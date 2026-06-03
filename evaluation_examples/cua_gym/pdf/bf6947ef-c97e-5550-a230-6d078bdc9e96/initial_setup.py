"""
Initial Setup: Multi-app PDF workflow - cover letter template + resume
Task ID: pdf_cross_137
Domain: pdf (with LibreOffice Writer)
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'pdf_cross_137'
DOCS_DIR = f'{WORKDIR}/Documents'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_cover_letter_odt():
    """Create cover_letter.odt with [APPLICANT_NAME] and [POSITION] placeholders."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    odt_path = f'{DOCS_DIR}/cover_letter.odt'

    doc = Document()

    # Set narrow margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Sender header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = header_para.add_run("[APPLICANT_NAME]")
    run.font.size = Pt(12)
    run.font.name = "Calibri"

    addr_para = doc.add_paragraph()
    addr_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run2 = addr_para.add_run("4821 Maple Ridge Drive\nSeattle, WA 98101\napplicant@email.com\n(206) 555-0183")
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run("March 25, 2026")
    date_run.font.size = Pt(11)
    date_run.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Recipient
    recip_para = doc.add_paragraph()
    recip_run = recip_para.add_run(
        "Hiring Manager\nTechVision Solutions Inc.\n1000 Innovation Way\nSan Francisco, CA 94105"
    )
    recip_run.font.size = Pt(11)
    recip_run.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Subject line
    subj_para = doc.add_paragraph()
    subj_run = subj_para.add_run("Re: Application for [POSITION]")
    subj_run.font.size = Pt(11)
    subj_run.bold = True
    subj_run.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Salutation
    sal_para = doc.add_paragraph()
    sal_run = sal_para.add_run("Dear Hiring Manager,")
    sal_run.font.size = Pt(11)
    sal_run.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Body paragraph 1
    p1 = doc.add_paragraph()
    r1 = p1.add_run(
        "I am writing to express my strong interest in the [POSITION] role at TechVision "
        "Solutions Inc. With over eight years of progressive experience in software engineering "
        "and a proven record of delivering scalable, high-impact systems, I am confident that my "
        "background aligns closely with what your team is seeking."
    )
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Body paragraph 2
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "In my most recent position at Nexus Technologies, I led the architecture and deployment of "
        "a distributed microservices platform that reduced system latency by 42% and improved overall "
        "reliability to 99.97% uptime. I collaborated closely with cross-functional teams including "
        "product management, QA, and DevOps to ensure timely delivery of quarterly milestones. My "
        "expertise spans Python, Go, Kubernetes, and cloud-native infrastructure on both AWS and GCP."
    )
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Body paragraph 3
    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        "I am particularly excited about TechVision's mission to democratize data intelligence and "
        "your recent work on real-time analytics pipelines. I believe my experience building "
        "event-driven architectures and mentoring junior engineers would add immediate value to "
        "your growing engineering organization."
    )
    r3.font.size = Pt(11)
    r3.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Closing
    close_para = doc.add_paragraph()
    close_run = close_para.add_run(
        "Thank you for your time and consideration. I look forward to the opportunity to discuss "
        "how my skills and passion can contribute to TechVision's continued success. Please find "
        "my resume attached for your review."
    )
    close_run.font.size = Pt(11)
    close_run.font.name = "Calibri"

    doc.add_paragraph()  # blank line

    # Sign-off
    sign_para = doc.add_paragraph()
    sign_run = sign_para.add_run("Sincerely,")
    sign_run.font.size = Pt(11)
    sign_run.font.name = "Calibri"

    doc.add_paragraph()  # blank space for signature
    doc.add_paragraph()

    name_para = doc.add_paragraph()
    name_run = name_para.add_run("[APPLICANT_NAME]")
    name_run.font.size = Pt(11)
    name_run.font.name = "Calibri"
    name_run.bold = True

    doc.save(odt_path)
    print(f'cover_letter.odt created: {odt_path}')


def create_resume_pdf():
    """Create resume.pdf - a realistic 2-page resume."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import black, HexColor, white
    from reportlab.lib.units import inch

    pdf_path = f'{DOCS_DIR}/resume.pdf'

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Title'],
        fontSize=22,
        textColor=HexColor("#1A3A5C"),
        spaceAfter=4,
        fontName='Helvetica-Bold',
    )
    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor("#444444"),
        spaceAfter=8,
        alignment=1,  # center
    )
    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Normal'],
        fontSize=12,
        fontName='Helvetica-Bold',
        textColor=HexColor("#1A3A5C"),
        spaceBefore=12,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=black,
        spaceAfter=4,
        leading=14,
    )
    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=black,
        spaceAfter=3,
        leftIndent=18,
        leading=13,
    )
    job_title_style = ParagraphStyle(
        'JobTitle',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceAfter=2,
    )
    company_style = ParagraphStyle(
        'CompanyStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor("#555555"),
        spaceAfter=4,
        fontName='Helvetica-Oblique',
    )

    story = []

    # --- PAGE 1 ---

    # Name
    story.append(Paragraph("[APPLICANT_NAME]", name_style))

    # Contact
    story.append(Paragraph(
        "4821 Maple Ridge Drive, Seattle, WA 98101  |  (206) 555-0183  |  applicant@email.com  |  linkedin.com/in/applicant",
        contact_style
    ))
    story.append(HRFlowable(width="100%", thickness=1.5, color=HexColor("#1A3A5C")))
    story.append(Spacer(1, 8))

    # Summary
    story.append(Paragraph("PROFESSIONAL SUMMARY", section_header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#AAAAAA")))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "Results-driven software engineering professional with 8+ years of experience designing and "
        "implementing robust, cloud-native systems. Demonstrated expertise in distributed systems, "
        "microservices architecture, and DevOps practices. Strong communicator with a track record of "
        "leading cross-functional teams and delivering projects on time and within budget.",
        body_style
    ))

    # Experience
    story.append(Spacer(1, 6))
    story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#AAAAAA")))
    story.append(Spacer(1, 4))

    # Job 1
    exp_table_data = [[
        Paragraph("<b>Nexus Technologies — Seattle, WA</b>", body_style),
        Paragraph("<i>Jan 2021 – Present</i>", ParagraphStyle('DateStyle', parent=styles['Normal'], fontSize=10, alignment=2))
    ]]
    exp_table = Table(exp_table_data, colWidths=[4.5 * inch, 2.5 * inch])
    exp_table.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
    story.append(exp_table)

    story.append(Paragraph("Staff Software Engineer", job_title_style))
    for bullet in [
        "• Architected distributed microservices platform serving 2M+ daily active users, achieving 99.97% uptime",
        "• Reduced end-to-end API latency by 42% through caching strategy redesign and async processing pipelines",
        "• Led team of 7 engineers across three scrum teams; conducted weekly 1:1s and quarterly performance reviews",
        "• Championed migration from monolith to event-driven architecture using Kafka and Kubernetes on GKE",
        "• Authored internal engineering playbooks adopted by 40+ engineers across the organization",
    ]:
        story.append(Paragraph(bullet, bullet_style))

    story.append(Spacer(1, 6))

    # Job 2
    exp_table_data2 = [[
        Paragraph("<b>DataBridge Corp — Portland, OR</b>", body_style),
        Paragraph("<i>Jun 2018 – Dec 2020</i>", ParagraphStyle('DateStyle2', parent=styles['Normal'], fontSize=10, alignment=2))
    ]]
    exp_table2 = Table(exp_table_data2, colWidths=[4.5 * inch, 2.5 * inch])
    exp_table2.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
    story.append(exp_table2)

    story.append(Paragraph("Senior Software Engineer", job_title_style))
    for bullet in [
        "• Designed and shipped real-time data ingestion service processing 50,000+ events per second",
        "• Integrated third-party analytics APIs (Segment, Mixpanel, Snowflake) to consolidate customer data",
        "• Reduced infrastructure costs by 28% through right-sizing EC2 instances and implementing spot fleet",
        "• Mentored three junior engineers; two received promotions within 18 months",
    ]:
        story.append(Paragraph(bullet, bullet_style))

    story.append(Spacer(1, 6))

    # Job 3
    exp_table_data3 = [[
        Paragraph("<b>Clearpath Systems — Austin, TX</b>", body_style),
        Paragraph("<i>Aug 2016 – May 2018</i>", ParagraphStyle('DateStyle3', parent=styles['Normal'], fontSize=10, alignment=2))
    ]]
    exp_table3 = Table(exp_table_data3, colWidths=[4.5 * inch, 2.5 * inch])
    exp_table3.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
    story.append(exp_table3)

    story.append(Paragraph("Software Engineer", job_title_style))
    for bullet in [
        "• Developed backend REST APIs in Python/Django serving mobile and web clients",
        "• Built automated regression testing suite reducing QA cycle time from 4 days to 6 hours",
        "• Collaborated with UX designers to implement responsive frontend features using React",
    ]:
        story.append(Paragraph(bullet, bullet_style))

    # --- PAGE 2 ---
    story.append(PageBreak())

    # Education
    story.append(Paragraph("EDUCATION", section_header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#AAAAAA")))
    story.append(Spacer(1, 4))

    edu_table_data = [[
        Paragraph("<b>University of Washington — Seattle, WA</b>", body_style),
        Paragraph("<i>Graduated May 2016</i>", ParagraphStyle('DateStyle4', parent=styles['Normal'], fontSize=10, alignment=2))
    ]]
    edu_table = Table(edu_table_data, colWidths=[4.5 * inch, 2.5 * inch])
    edu_table.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
    story.append(edu_table)
    story.append(Paragraph("Bachelor of Science in Computer Science", job_title_style))
    story.append(Paragraph("GPA: 3.78 / 4.00  |  Magna Cum Laude  |  Dean's List (all semesters)", body_style))

    story.append(Spacer(1, 10))

    # Technical Skills
    story.append(Paragraph("TECHNICAL SKILLS", section_header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#AAAAAA")))
    story.append(Spacer(1, 4))

    skills_data = [
        ["Languages:", "Python, Go, TypeScript, Java, Bash, SQL"],
        ["Cloud / Infra:", "AWS (EC2, Lambda, RDS, S3), GCP (GKE, BigQuery), Terraform, Helm"],
        ["Frameworks:", "FastAPI, Django, gRPC, React, Spring Boot"],
        ["Databases:", "PostgreSQL, Redis, MongoDB, Cassandra, Snowflake"],
        ["DevOps / Tools:", "Docker, Kubernetes, GitHub Actions, Jenkins, Datadog, PagerDuty"],
        ["Architecture:", "Microservices, Event-Driven (Kafka), REST, GraphQL, CQRS"],
    ]
    for row in skills_data:
        skill_para = Paragraph(f"<b>{row[0]}</b> {row[1]}", body_style)
        story.append(skill_para)

    story.append(Spacer(1, 10))

    # Certifications
    story.append(Paragraph("CERTIFICATIONS & AWARDS", section_header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#AAAAAA")))
    story.append(Spacer(1, 4))
    for cert in [
        "• AWS Certified Solutions Architect – Professional (2024)",
        "• Google Cloud Professional Cloud Architect (2023)",
        "• Certified Kubernetes Administrator (CKA) – CNCF (2022)",
        "• Nexus Technologies Engineering Excellence Award (2023)",
        "• DataBridge Corp 'Above and Beyond' Recognition (2020)",
    ]:
        story.append(Paragraph(cert, bullet_style))

    story.append(Spacer(1, 10))

    # Projects
    story.append(Paragraph("SELECTED PROJECTS", section_header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#AAAAAA")))
    story.append(Spacer(1, 4))

    story.append(Paragraph("<b>OpenMetrics Gateway</b> — Open Source (github.com/applicant/openmetrics-gw)", body_style))
    story.append(Paragraph(
        "• Built a high-throughput metrics aggregation gateway (10K+ GitHub stars) used in production by 200+ companies",
        bullet_style
    ))
    story.append(Spacer(1, 4))

    story.append(Paragraph("<b>StreamSync Engine</b> — Personal Project", body_style))
    story.append(Paragraph(
        "• Developed a real-time data synchronization engine supporting multi-region replication with <50ms lag",
        bullet_style
    ))

    story.append(Spacer(1, 10))

    # References
    story.append(Paragraph("REFERENCES", section_header_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#AAAAAA")))
    story.append(Spacer(1, 4))
    story.append(Paragraph("Available upon request.", body_style))

    doc.build(story)
    print(f'resume.pdf created: {pdf_path}')


def create_initial():
    os.makedirs(DOCS_DIR, exist_ok=True)

    create_cover_letter_odt()
    create_resume_pdf()

    print('All initial files created.')

    # GUI: Open LibreOffice Writer with the cover letter template
    launch_gui(f'libreoffice --writer "{DOCS_DIR}/cover_letter.odt"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with cover_letter.odt on DISPLAY=:0')


create_initial()
