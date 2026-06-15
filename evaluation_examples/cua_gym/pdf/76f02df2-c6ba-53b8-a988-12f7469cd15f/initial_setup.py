"""
Initial Setup: Multi-app workflow - GIMP letterhead + Writer proposal document
Task ID: pdf_cross_141
Domain: pdf (multi-app: GIMP, LibreOffice Writer, pymupdf)

Creates:
  - /home/user/letterhead.png  : 2000x250px blue gradient letterhead image
  - /home/user/proposal_draft.docx : Writer doc with 3 paragraphs, NO letterhead header, NO watermark
Opens the .docx in LibreOffice Writer for the agent to work with.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'

# ---------------------------------------------------------------------------
# GUI launch helper
# ---------------------------------------------------------------------------

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# ---------------------------------------------------------------------------
# Step 1: Create letterhead image with PIL
# ---------------------------------------------------------------------------

def create_letterhead():
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        subprocess.run(["pip3", "install", "Pillow"], check=True)
        from PIL import Image, ImageDraw, ImageFont

    width, height = 2000, 250

    # Create blue gradient (left: dark blue #003399, right: medium blue #0066CC)
    img = Image.new("RGB", (width, height), (0, 0, 0))
    draw = ImageDraw.Draw(img)

    # Draw gradient horizontally
    for x in range(width):
        ratio = x / width
        r = int(0 * (1 - ratio) + 0 * ratio)
        g = int(51 * (1 - ratio) + 102 * ratio)
        b = int(153 * (1 - ratio) + 204 * ratio)
        draw.line([(x, 0), (x, height)], fill=(r, g, b))

    # Add white company name text
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    fallback_paths = [
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
        "/usr/share/fonts/truetype/ubuntu/Ubuntu-B.ttf",
    ]
    font = None
    try:
        font = ImageFont.truetype(font_path, 100)
    except Exception:
        for fp in fallback_paths:
            try:
                font = ImageFont.truetype(fp, 100)
                break
            except Exception:
                continue
    if font is None:
        font = ImageFont.load_default()

    text = "Atlas Dynamics"
    # Calculate text bounding box for centering
    bbox = draw.textbbox((0, 0), text, font=font)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = (width - text_w) // 2
    y = (height - text_h) // 2 - bbox[1]  # offset for baseline

    draw.text((x, y), text, fill=(255, 255, 255), font=font)

    output_path = os.path.join(WORKDIR, "letterhead.png")
    img.save(output_path)
    print(f"Letterhead created: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Step 2: Create proposal draft .docx (NO letterhead, NO watermark, NOT as PDF)
# ---------------------------------------------------------------------------

def create_proposal_draft():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        subprocess.run(["pip3", "install", "python-docx"], check=True)
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Business Proposal: Strategic Partnership Initiative", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Prepared by / date line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prepared by: Atlas Dynamics — March 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()  # spacer

    # Section heading
    doc.add_heading("Executive Summary", level=2)

    # Paragraph 1
    p1 = doc.add_paragraph(
        "Atlas Dynamics is pleased to present this strategic partnership proposal to Nexus "
        "Innovations Group. Over the past decade, Atlas Dynamics has established itself as a "
        "leader in enterprise automation solutions, delivering measurable efficiency gains to "
        "over 350 clients across North America, Europe, and Asia-Pacific. This proposal outlines "
        "a collaborative framework designed to accelerate product development cycles, expand "
        "market reach, and generate substantial value for both organizations. Our combined "
        "capabilities in AI-driven workflow optimization and Nexus's robust distribution network "
        "represent a compelling opportunity for shared growth."
    )
    p1.style = doc.styles['Normal']

    doc.add_paragraph()  # spacer

    # Section heading
    doc.add_heading("Proposed Collaboration Scope", level=2)

    # Paragraph 2
    p2 = doc.add_paragraph(
        "The proposed partnership encompasses three primary areas of collaboration. First, joint "
        "product development: Atlas Dynamics will contribute its proprietary automation engine "
        "while Nexus Innovations provides domain expertise in logistics and supply-chain "
        "management. Together, we will co-develop an integrated platform capable of handling "
        "end-to-end procurement workflows with minimal human intervention. Second, co-marketing "
        "initiatives: both organizations will align marketing calendars to launch joint campaigns "
        "targeting Fortune 500 enterprises in Q3 2026. Third, shared revenue streams: a "
        "transparent revenue-sharing model will be established, with Atlas Dynamics retaining "
        "60% of software licensing fees and Nexus receiving 40% in exchange for exclusive "
        "distribution rights in designated territories."
    )
    p2.style = doc.styles['Normal']

    doc.add_paragraph()  # spacer

    # Section heading
    doc.add_heading("Financial Projections and Next Steps", level=2)

    # Paragraph 3
    p3 = doc.add_paragraph(
        "Based on current market analysis and conservative adoption forecasts, the partnership "
        "is projected to generate combined revenues of $12.4 million in Year 1, scaling to "
        "$31.7 million by Year 3. Atlas Dynamics commits to investing $2.1 million in dedicated "
        "R&D resources over the first 18 months to ensure timely delivery of the co-developed "
        "platform. As an immediate next step, we propose scheduling a formal due-diligence "
        "workshop in April 2026, bringing together senior technical and commercial teams from "
        "both organizations. Upon mutual agreement on the commercial terms outlined herein, a "
        "definitive partnership agreement will be drafted and executed no later than June 2026. "
        "We look forward to building a long-term, mutually beneficial relationship with Nexus "
        "Innovations Group and are confident that this collaboration will position both companies "
        "as market leaders in the enterprise automation sector."
    )
    p3.style = doc.styles['Normal']

    output_path = os.path.join(WORKDIR, "proposal_draft.docx")
    doc.save(output_path)
    print(f"Proposal draft created: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    create_letterhead()
    docx_path = create_proposal_draft()

    # GUI: open the draft in LibreOffice Writer for the agent
    launch_gui(f'libreoffice --writer "{docx_path}"', delay_sec=3.0)
    print("GUI_READY: launched LibreOffice Writer with proposal_draft.docx on DISPLAY=:0")


main()
