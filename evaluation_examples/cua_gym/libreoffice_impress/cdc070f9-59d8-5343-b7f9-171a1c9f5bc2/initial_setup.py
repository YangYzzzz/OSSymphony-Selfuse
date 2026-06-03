"""
Initial Setup: Create Strategy_Outline.docx and open LibreOffice Impress
Task ID: impress_wf_035
Domain: libreoffice_impress
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

WORKDIR = '/home/user'
DESKTOP = f'{WORKDIR}/Desktop'
TASK_ID = 'impress_wf_035'
OUTPUT_DOC = f'{DESKTOP}/Strategy_Outline.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# --- Strategic pillars and initiatives data ---
PILLARS = [
    {
        "name": "Market Expansion",
        "initiatives": [
            {
                "title": "Asia-Pacific Regional Entry",
                "body": (
                    "Establish regional offices in Singapore and Tokyo to penetrate "
                    "high-growth markets across Southeast Asia and Japan. Target initial "
                    "revenue contribution of $12M within the first 18 months through "
                    "strategic partnerships with local distributors and channel partners."
                ),
            },
            {
                "title": "Mid-Market Segment Penetration",
                "body": (
                    "Develop tailored product bundles and pricing tiers for mid-market "
                    "companies with 200-2000 employees. Launch a dedicated sales team "
                    "focused on vertical-specific messaging for healthcare, manufacturing, "
                    "and financial services sectors."
                ),
            },
            {
                "title": "Strategic Acquisition Pipeline",
                "body": (
                    "Identify and evaluate 3-5 acquisition targets in complementary "
                    "technology segments. Prioritize companies with annual recurring "
                    "revenue between $5M and $20M that would accelerate our platform "
                    "capabilities in data analytics and AI-driven insights."
                ),
            },
        ],
    },
    {
        "name": "Digital Transformation",
        "initiatives": [
            {
                "title": "Cloud Infrastructure Migration",
                "body": (
                    "Migrate 80% of on-premise workloads to a multi-cloud architecture "
                    "using AWS and Azure by Q3 2026. This initiative will reduce "
                    "infrastructure costs by an estimated 35% while improving system "
                    "uptime from 99.5% to 99.95%."
                ),
            },
            {
                "title": "Customer Experience Platform",
                "body": (
                    "Deploy an integrated digital experience platform that unifies "
                    "web, mobile, and in-app touchpoints. Implement real-time "
                    "personalization engines powered by machine learning to increase "
                    "customer engagement scores by 40%."
                ),
            },
            {
                "title": "Data-Driven Decision Framework",
                "body": (
                    "Build a centralized data lake and self-service analytics portal "
                    "accessible to all business units. Establish data governance "
                    "policies and train 150+ employees on advanced analytics tools "
                    "including Tableau and Power BI."
                ),
            },
        ],
    },
    {
        "name": "Talent Development",
        "initiatives": [
            {
                "title": "Leadership Academy Program",
                "body": (
                    "Create a 12-month executive development program for high-potential "
                    "managers. Partner with Wharton and INSEAD for quarterly immersion "
                    "modules covering strategic thinking, change management, and "
                    "cross-cultural leadership."
                ),
            },
            {
                "title": "Technical Skills Upskilling",
                "body": (
                    "Launch a comprehensive learning platform offering 200+ courses "
                    "in cloud computing, AI/ML, cybersecurity, and DevOps. Allocate "
                    "$2.5M annually for certification reimbursements and dedicate "
                    "10% of work hours to structured learning."
                ),
            },
            {
                "title": "Diversity and Inclusion Initiative",
                "body": (
                    "Achieve 45% gender diversity in leadership roles by 2027 through "
                    "targeted recruitment, mentorship programs, and inclusive promotion "
                    "criteria. Establish Employee Resource Groups for underrepresented "
                    "communities and conduct quarterly inclusion surveys."
                ),
            },
        ],
    },
    {
        "name": "Operational Excellence",
        "initiatives": [
            {
                "title": "Process Automation at Scale",
                "body": (
                    "Deploy robotic process automation across finance, HR, and supply "
                    "chain operations. Target 60% reduction in manual processing time "
                    "for invoice management, employee onboarding, and inventory "
                    "reconciliation workflows."
                ),
            },
            {
                "title": "Supply Chain Optimization",
                "body": (
                    "Implement predictive analytics for demand forecasting and "
                    "inventory optimization. Reduce average lead times by 25% and "
                    "carrying costs by $8M through strategic vendor consolidation "
                    "and just-in-time procurement practices."
                ),
            },
            {
                "title": "Quality Management System Upgrade",
                "body": (
                    "Transition to ISO 9001:2025 standards with integrated digital "
                    "quality management. Implement automated testing and inspection "
                    "protocols that reduce defect rates by 50% and customer complaints "
                    "by 30% within 12 months."
                ),
            },
        ],
    },
    {
        "name": "Innovation and R&D",
        "initiatives": [
            {
                "title": "AI Research Lab Establishment",
                "body": (
                    "Open a dedicated AI research facility in Austin, TX with an "
                    "initial team of 25 researchers and engineers. Focus on "
                    "generative AI applications, natural language processing, and "
                    "computer vision solutions for enterprise customers."
                ),
            },
            {
                "title": "Innovation Incubator Program",
                "body": (
                    "Allocate $5M seed funding for an internal incubator that supports "
                    "employee-driven innovation projects. Run biannual pitch "
                    "competitions and provide winning teams with 6-month dedicated "
                    "development sprints and go-to-market support."
                ),
            },
            {
                "title": "University Partnership Network",
                "body": (
                    "Establish research partnerships with MIT, Stanford, and ETH "
                    "Zurich to co-develop next-generation technologies. Fund 10 "
                    "PhD fellowships annually and create a joint intellectual "
                    "property framework for commercializing research outcomes."
                ),
            },
        ],
    },
]


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Title
    title_para = doc.add_heading("Strategic Growth Plan 2025-2027", level=0)

    # Intro paragraph
    doc.add_paragraph(
        "This document outlines the five strategic pillars that will guide our "
        "organization's growth trajectory over the next three fiscal years. Each "
        "pillar is supported by three actionable initiatives with measurable "
        "outcomes and defined timelines."
    )

    # Add each pillar as Heading 1, initiatives as Heading 2
    for pillar in PILLARS:
        doc.add_heading(pillar["name"], level=1)
        for initiative in pillar["initiatives"]:
            doc.add_heading(initiative["title"], level=2)
            doc.add_paragraph(initiative["body"])

    doc.save(OUTPUT_DOC)
    print(f"Initial document created: {OUTPUT_DOC}")

    # Open LibreOffice Impress (blank) for GUI-ready state
    launch_gui('libreoffice --impress', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Impress with DISPLAY=:0")


create_initial()
