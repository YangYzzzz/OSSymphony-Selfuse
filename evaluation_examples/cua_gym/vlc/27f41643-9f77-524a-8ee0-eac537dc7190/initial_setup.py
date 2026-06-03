"""
Initial Setup: Navigate to timestamp 1:45 in VLC tutorial video, take a snapshot,
then insert that snapshot into LibreOffice Writer document.
Task ID: vlcplay_013
Domain: vlc
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'vlcplay_013'
VIDEO_DIR = f'{WORKDIR}/Videos'
DOC_DIR = f'{WORKDIR}/Documents'
PICTURES_DIR = f'{WORKDIR}/Pictures'
VIDEO_PATH = f'{VIDEO_DIR}/tutorial.avi'
DOC_PATH = f'{DOC_DIR}/notes.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    env["VLC_VERBOSE"] = "-1"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_video():
    """Create a tutorial-style test video (~2 minutes) with distinct visual content at different timestamps."""
    os.makedirs(VIDEO_DIR, exist_ok=True)

    # Use a shell command string to handle ffmpeg filter escaping correctly
    cmd = (
        'ffmpeg -y '
        '-f lavfi -i "smptebars=duration=120:size=1280x720:rate=24" '
        '-f lavfi -i "anullsrc=r=44100:cl=stereo" '
        '-c:v libx264 -preset ultrafast -pix_fmt yuv420p '
        '-c:a aac -shortest '
        f'"{VIDEO_PATH}"'
    )
    result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
    if result.returncode != 0:
        # Fallback: use testsrc if smptebars is not available
        cmd_fallback = (
            'ffmpeg -y '
            '-f lavfi -i "testsrc=duration=120:size=1280x720:rate=24" '
            '-f lavfi -i "anullsrc=r=44100:cl=stereo" '
            '-c:v libx264 -preset ultrafast -pix_fmt yuv420p '
            '-c:a aac -shortest '
            f'"{VIDEO_PATH}"'
        )
        subprocess.run(cmd_fallback, shell=True, check=True, capture_output=True)
    print(f"Video created: {VIDEO_PATH}")


def create_document():
    """Create a notes document with realistic text content but NO images."""
    os.makedirs(DOC_DIR, exist_ok=True)

    # Use python-docx to create the document
    subprocess.run([
        "python3", "-c", """
import sys
sys.path.insert(0, '/usr/lib/python3/dist-packages')
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Python Data Analysis Tutorial Notes', level=1)

# Section 1
doc.add_heading('1. Setting Up the Environment', level=2)
doc.add_paragraph(
    'Before starting, ensure you have Python 3.10 or later installed. '
    'We will use pandas, numpy, and matplotlib for this tutorial. '
    'Install them using pip:'
)
doc.add_paragraph('pip install pandas numpy matplotlib seaborn', style='Intense Quote')
doc.add_paragraph(
    'Once installed, open a Jupyter notebook or your preferred IDE. '
    'We recommend VS Code with the Python extension for the best experience.'
)

# Section 2
doc.add_heading('2. Loading and Inspecting Data', level=2)
doc.add_paragraph(
    'The first step in any data analysis project is loading your dataset. '
    'We will work with a sales dataset containing quarterly revenue figures '
    'for a mid-sized technology company from 2020 to 2024.'
)
doc.add_paragraph(
    'Key columns in our dataset: Date, Product, Region, Revenue, Units_Sold, '
    'Customer_Segment. The dataset contains 2,847 records across 5 product lines.'
)

# Section 3
doc.add_heading('3. Data Cleaning', level=2)
doc.add_paragraph(
    'Missing values were found in the Revenue column (23 records) and '
    'the Customer_Segment column (8 records). We applied median imputation '
    'for Revenue and mode imputation for Customer_Segment.'
)
doc.add_paragraph(
    'Outliers were identified using the IQR method. Three records with '
    'Revenue > $500,000 were flagged as data entry errors and corrected '
    'by cross-referencing the original invoices.'
)

# Section 4
doc.add_heading('4. Visualization Screenshots', level=2)
doc.add_paragraph(
    'Below this section, insert relevant screenshots and visualizations '
    'from the tutorial video to supplement these notes.'
)
doc.add_paragraph(
    '[Insert VLC snapshot here]'
)

# Section 5
doc.add_heading('5. Key Findings', level=2)
doc.add_paragraph(
    'Revenue increased by 34% year-over-year from 2022 to 2023. '
    'The Enterprise segment showed the strongest growth at 47%, '
    'while the SMB segment remained flat. Regional analysis reveals '
    'that APAC overtook EMEA as the second-largest market in Q3 2023.'
)

doc.save('/home/user/Documents/notes.docx')
print('Document created: /home/user/Documents/notes.docx')
"""
    ], check=True, capture_output=True)


def create_pictures_dir():
    """Ensure Pictures directory exists for VLC snapshot output."""
    os.makedirs(PICTURES_DIR, exist_ok=True)
    print(f"Pictures directory ready: {PICTURES_DIR}")


def main():
    # Create the video file
    create_video()

    # Create the Writer document (no images)
    create_document()

    # Ensure Pictures directory exists
    create_pictures_dir()

    # Launch VLC with the tutorial video at the beginning
    launch_gui(f'vlc --no-video-title-show "{VIDEO_PATH}"', delay_sec=3.0)
    print('GUI_READY: VLC launched with tutorial.avi')

    # Launch LibreOffice Writer with the notes document
    launch_gui(f'libreoffice --writer "{DOC_PATH}"', delay_sec=3.0)
    print('GUI_READY: LibreOffice Writer launched with notes.docx')


main()
