"""
FINAL REWARD SCRIPT - SUCCESS
Task: LibreOffice decided to wrap my data in something labeled “Table 1”, but I actually need it as plain text—each original cell separated by a single comma (,). How do I quickly flip Table 1 back to that comma-delimited line?
Generated: 2025-09-10 13:58:09
Status: success
Model: azure-o3
Total Steps: 4
"""

import os
from docx import Document


def verify_task(file_path: str) -> float:
    """
    Verify that the user successfully converted a LibreOffice/Word table back
    into a single comma-delimited text line.

    Scoring (progressive – max 1.0):
      • 0.5 pts  – ALL tables were removed from the document (table count == 0)
      • 0.5 pts  – A paragraph containing a comma-delimited line (≥2 commas)
                   exists in the document
    Returns a float between 0.0 and 1.0
    """
    print(f"Verifying file: {file_path}")

    # --- Basic safety checks (no points) -----------------------------------
    if not os.path.exists(file_path):
        print("✗ File not found – cannot verify task")
        return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to open DOCX: {e}")
        return 0.0

    total_score = 0.0  # progressive score

    # --- Requirement 1: Document must contain *no* tables -------------------
    table_count = len(doc.tables)
    print(f"Table count detected: {table_count}")
    if table_count == 0:
        print("✓ All tables removed (0.5 pts)")
        total_score += 0.5
    else:
        print("✗ Tables are still present (0 pts)")

    # --- Requirement 2: Comma-delimited line present ------------------------
    comma_line_found = False
    for para in doc.paragraphs:
        text = para.text.strip()
        # A valid comma-delimited line should contain at least two commas
        if text.count(',') >= 2:
            comma_line_found = True
            print(f"✓ Found comma-delimited line: '{text}' (0.5 pts)")
            break

    if comma_line_found:
        total_score += 0.5
    else:
        print("✗ No suitable comma-delimited line found (0 pts)")

    # -----------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"Final score: {final_score}")
    return final_score


if __name__ == "__main__":
    DOC_PATH = "/home/user/libreoffice_decided_to_wrap_my_data_in_something_labeled_table_1_but_i_actually_need_it_as_plain_tex.docx"
    reward = verify_task(DOC_PATH)
    print(f"REWARD: {reward}")

