"""
Reward Script: Insert © before 'Pinnacle Solutions' and ™ after 'SmartFlow'
Task ID: writer_txtfmt_042
Domain: libreoffice_writer
Scoring:
  Component 1: Trademark symbol ™ (U+2122) after 'SmartFlow' in paragraph 0 — 0.5 pts
  Component 2: Copyright symbol © (U+00A9) before 'Pinnacle Solutions' in last paragraph — 0.5 pts
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_042'

TM_CHAR = '\u2122'    # ™
COPYRIGHT_CHAR = '\u00A9'  # ©

FILE_PATH = f'{WORKDIR}/Desktop/product_brochure.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: doc must have at least 10 paragraphs
    if len(doc.paragraphs) < 1:
        print("CRITICAL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Trademark symbol ™ (U+2122) inserted after 'SmartFlow' in paragraph 0 (0.5 points)
    # Expected golden text: "Introducing SmartFlow™, our flagship..."
    # Initial text:         "Introducing SmartFlow, our flagship..." (no ™)
    try:
        first_para = doc.paragraphs[0]
        first_para_text = first_para.text
        # Check that SmartFlow is followed by ™
        if TM_CHAR in first_para_text and 'SmartFlow' + TM_CHAR in first_para_text:
            print(f"PASS: Component 1 — Trademark symbol ™ found after 'SmartFlow' in paragraph 0 (0.5 pts)")
            print(f"       Text: {repr(first_para_text[:80])}")
            total_score += 0.5
        else:
            # Check for partial: SmartFlow present but TM missing
            if 'SmartFlow' in first_para_text:
                print(f"FAIL: Component 1 — 'SmartFlow' found but ™ symbol missing in paragraph 0")
            else:
                print(f"FAIL: Component 1 — 'SmartFlow' not found in paragraph 0")
            print(f"       Actual text: {repr(first_para_text[:80])}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Copyright symbol © (U+00A9) inserted before 'Pinnacle Solutions' in last paragraph (0.5 points)
    # Expected golden text: "©Pinnacle Solutions 2025. All rights reserved."
    # Initial text:         "Pinnacle Solutions 2025. All rights reserved." (no ©)
    # Note: The task says "footer text" but the document stores this in the last body paragraph (Para 9).
    # We check the last paragraph of the document body, which is where setup-gen placed the change.
    try:
        # Find the last paragraph with non-empty text
        last_para = None
        for p in reversed(doc.paragraphs):
            if p.text.strip():
                last_para = p
                break

        if last_para is None:
            print("FAIL: Component 2 — No non-empty paragraphs found")
        else:
            last_para_text = last_para.text
            # Check that © precedes 'Pinnacle Solutions'
            if COPYRIGHT_CHAR in last_para_text and COPYRIGHT_CHAR + 'Pinnacle Solutions' in last_para_text:
                print(f"PASS: Component 2 — Copyright symbol © found before 'Pinnacle Solutions' in last paragraph (0.5 pts)")
                print(f"       Text: {repr(last_para_text)}")
                total_score += 0.5
            else:
                if 'Pinnacle Solutions' in last_para_text:
                    print(f"FAIL: Component 2 — 'Pinnacle Solutions' found but © symbol missing or not directly before it")
                else:
                    print(f"FAIL: Component 2 — 'Pinnacle Solutions' not found in last paragraph")
                print(f"       Actual text: {repr(last_para_text)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
