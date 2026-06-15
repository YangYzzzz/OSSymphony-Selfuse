"""
Initial Setup: Product brochure document with SmartFlow and Pinnacle Solutions references
Task ID: writer_txtfmt_042
Domain: libreoffice_writer

Creates product_brochure.docx at ~/Desktop/ with:
- First paragraph mentions 'SmartFlow' WITHOUT trademark symbol
- Last line reads 'Pinnacle Solutions 2025. All rights reserved.' WITHOUT copyright symbol
- All text in 12pt Calibri
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'product_brochure'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_style(run, font_name='Calibri', font_size_pt=12):
    """Set font name and size for a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Paragraph 1: Product introduction mentioning SmartFlow ---
    p1 = doc.add_paragraph()
    run1 = p1.add_run(
        'Introducing SmartFlow, our flagship workflow automation platform designed '
        'to streamline your business processes and maximize productivity across all departments.'
    )
    set_run_style(run1)

    # --- Paragraph 2: Features section ---
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        'SmartFlow delivers unmatched performance with its intuitive drag-and-drop interface, '
        'real-time analytics dashboard, and seamless integration with over 200 enterprise applications.'
    )
    set_run_style(run2)

    # --- Paragraph 3: Benefits ---
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        'Key Benefits:'
    )
    set_run_style(run3)
    run3.bold = True

    # --- Paragraph 4: Benefit list item 1 ---
    p4 = doc.add_paragraph()
    run4 = p4.add_run(
        '- Reduce operational costs by up to 40% through intelligent process automation'
    )
    set_run_style(run4)

    # --- Paragraph 5: Benefit list item 2 ---
    p5 = doc.add_paragraph()
    run5 = p5.add_run(
        '- Accelerate project delivery timelines with automated task routing and approval workflows'
    )
    set_run_style(run5)

    # --- Paragraph 6: Benefit list item 3 ---
    p6 = doc.add_paragraph()
    run6 = p6.add_run(
        '- Improve compliance and audit readiness with comprehensive activity logging'
    )
    set_run_style(run6)

    # --- Paragraph 7: Customer success ---
    p7 = doc.add_paragraph()
    run7 = p7.add_run(
        'Trusted by over 5,000 organizations worldwide, SmartFlow has transformed how leading '
        'enterprises manage their critical business workflows. From healthcare to finance, '
        'our clients report an average 60% reduction in manual processing time within the first 90 days.'
    )
    set_run_style(run7)

    # --- Paragraph 8: Pricing/Plans ---
    p8 = doc.add_paragraph()
    run8 = p8.add_run(
        'SmartFlow is available in three tiers: Starter ($49/month), Professional ($149/month), '
        'and Enterprise (custom pricing). All plans include 24/7 technical support, '
        'automatic updates, and a 30-day free trial.'
    )
    set_run_style(run8)

    # --- Paragraph 9: Contact information ---
    p9 = doc.add_paragraph()
    run9 = p9.add_run(
        'For more information or to schedule a demo, contact our sales team at '
        'sales@pinnaclesolutions.com or call 1-800-PINNACLE (1-800-746-6225).'
    )
    set_run_style(run9)

    # --- Paragraph 10: Footer/Copyright line (last line) ---
    # NOTE: NO copyright symbol before 'Pinnacle Solutions' (that is the task)
    p10 = doc.add_paragraph()
    run10 = p10.add_run(
        'Pinnacle Solutions 2025. All rights reserved.'
    )
    set_run_style(run10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
