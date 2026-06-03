"""
Reward Script: Merge first three cells in row 1 of table to create single title cell
Task ID: writer_tm_003
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.5): Row 0 first cell has gridSpan=3 (horizontal merge)
  - Component 2 (0.3): Merged cell text contains 'Q1 Budget'
  - Component 3 (0.2): Rows 1-4 remain unmerged (no gridSpan > 1)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_003'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that row 0 of the table has been merged into a single cell
    spanning all 3 columns, while rows 1-4 remain unmerged.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table must have at least 5 rows
    if len(table.rows) < 5:
        print(f"FAIL: Table has {len(table.rows)} rows, expected at least 5")
        print("REWARD: 0.0")
        return 0.0

    # Helper: check if row 0 first cell has gridSpan=3
    merge_detected = False
    try:
        row0 = table.rows[0]
        first_tc = row0.cells[0]._tc
        tcPr = first_tc.find(qn('w:tcPr'))
        grid_span_val = None
        if tcPr is not None:
            gs = tcPr.find(qn('w:gridSpan'))
            if gs is not None:
                grid_span_val = gs.get(qn('w:val'))
        if grid_span_val is not None and int(grid_span_val) == 3:
            merge_detected = True
    except Exception:
        pass

    # Component 1: Row 0 first cell has gridSpan=3 (0.5 points)
    # This is the core task: merge 3 cells into one spanning cell
    try:
        if merge_detected:
            print(f"PASS: Component 1 — Row 0 cell(0,0) has gridSpan=3 (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Expected gridSpan=3 in row 0, found gridSpan={grid_span_val}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Merged cell contains 'Q1 Budget' AND merge exists (0.3 points)
    # Both the merge AND text preservation are required — anchored to the task change
    try:
        row0 = table.rows[0]
        merged_text = row0.cells[0].text.strip()
        if merge_detected and 'Q1 Budget' in merged_text:
            print(f"PASS: Component 2 — Merged cell contains 'Q1 Budget' (text: {repr(merged_text)}) (0.3 pts)")
            total_score += 0.3
        elif not merge_detected:
            print(f"FAIL: Component 2 — Merge not detected, so text check not applicable")
        else:
            print(f"FAIL: Component 2 — Expected 'Q1 Budget' in merged cell, found: {repr(merged_text)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Row 0 is merged AND rows 1-4 remain unmerged (0.2 points)
    # Verifies merge is correctly scoped to only row 0 — anchored to the task change
    try:
        all_unmerged = True
        for ri in range(1, min(5, len(table.rows))):
            row = table.rows[ri]
            for ci, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    gs = tcPr.find(qn('w:gridSpan'))
                    if gs is not None:
                        val = int(gs.get(qn('w:val'), '1'))
                        if val > 1:
                            all_unmerged = False
                            print(f"FAIL: Component 3 — Row {ri} cell {ci} has gridSpan={val}, expected no merge")
                            break
            if not all_unmerged:
                break

        if merge_detected and all_unmerged:
            print(f"PASS: Component 3 — Row 0 merged AND rows 1-4 remain unmerged (0.2 pts)")
            total_score += 0.2
        elif not merge_detected:
            print(f"FAIL: Component 3 — Merge not detected in row 0, so scope check not applicable")
        elif not all_unmerged:
            pass  # already printed above
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist unsaved state before verification
persist_app_state("libreoffice_writer")

# Run verification
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
