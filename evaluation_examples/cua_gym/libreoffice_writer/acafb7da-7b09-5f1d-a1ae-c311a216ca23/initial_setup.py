"""
Initial Setup: Create a Budget Report document with a 3x5 table.
Task ID: writer_tm_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a heading
    heading = doc.add_heading("Annual Budget Report", level=1)

    # Add an introductory paragraph
    intro = doc.add_paragraph(
        "The following table summarizes our quarterly budget allocation "
        "across key departments. Please review the figures and provide "
        "feedback before the next planning meeting on April 15th."
    )

    # Create a 5-row x 3-column table
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"

    # Row 0 (header row): "Q1 Budget" in cell A1, B1 and C1 empty
    # Cells are NOT merged in initial state
    cell_a1 = table.cell(0, 0)
    cell_a1.text = "Q1 Budget"
    # Make header bold
    for run in cell_a1.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(12)

    # B1 and C1 are left empty (as specified)
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = ""

    # Row 1 (data row): column headers
    table.cell(1, 0).text = "Department"
    table.cell(1, 1).text = "Allocated ($)"
    table.cell(1, 2).text = "Spent ($)"
    # Bold the column headers
    for col_idx in range(3):
        for run in table.cell(1, col_idx).paragraphs[0].runs:
            run.bold = True

    # Row 2: Engineering
    table.cell(2, 0).text = "Engineering"
    table.cell(2, 1).text = "125,000"
    table.cell(2, 2).text = "98,750"

    # Row 3: Marketing
    table.cell(3, 0).text = "Marketing"
    table.cell(3, 1).text = "85,000"
    table.cell(3, 2).text = "72,340"

    # Row 4: Operations
    table.cell(4, 0).text = "Operations"
    table.cell(4, 1).text = "67,500"
    table.cell(4, 2).text = "54,210"

    # Add a closing paragraph
    doc.add_paragraph(
        "Note: All amounts are in USD. Figures are subject to revision "
        "pending final audit results from the finance team."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
