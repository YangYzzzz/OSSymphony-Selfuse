"""
Initial Setup: Create invoice document with 4x10 table, INVOICE in A1 left-aligned 12pt regular
Task ID: writer_tm_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_020'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a heading for context
    heading = doc.add_heading('Company Invoice', level=1)

    # Create 4x10 table
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Table Grid'

    # Row 0 (Row 1 in task terms): "INVOICE" in cell A1, B1-D1 empty
    cell_a1 = table.cell(0, 0)
    cell_a1.paragraphs[0].clear()
    run = cell_a1.paragraphs[0].add_run('INVOICE')
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.name = 'Calibri'
    # Left-aligned (default, but set explicitly)
    cell_a1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # B1, C1, D1 are empty (default)

    # Row 1 (index 1): Column headers
    headers = ['Item', 'Quantity', 'Unit Price', 'Total']
    for col_idx, h in enumerate(headers):
        cell = table.cell(1, col_idx)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # Rows 2-9 (index 2-9): Invoice line items with realistic data
    invoice_data = [
        ['Website Redesign',        '1',  '$4,500.00',  '$4,500.00'],
        ['Logo Design Package',     '2',  '$1,200.00',  '$2,400.00'],
        ['SEO Optimization',        '1',  '$2,800.00',  '$2,800.00'],
        ['Social Media Setup',      '3',  '$750.00',    '$2,250.00'],
        ['Content Writing (pages)', '10', '$350.00',    '$3,500.00'],
        ['Photography Session',     '1',  '$1,600.00',  '$1,600.00'],
        ['Email Template Design',   '4',  '$450.00',    '$1,800.00'],
        ['Hosting Setup (annual)',   '1',  '$1,100.00',  '$1,100.00'],
    ]

    for row_idx, row_data in enumerate(invoice_data, start=2):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # Add a paragraph after the table for context
    doc.add_paragraph('')
    summary = doc.add_paragraph()
    run = summary.add_run('Subtotal: $19,950.00')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    tax_para = doc.add_paragraph()
    run = tax_para.add_run('Tax (8.5%): $1,695.75')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    total_para = doc.add_paragraph()
    run = total_para.add_run('Grand Total: $21,645.75')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
