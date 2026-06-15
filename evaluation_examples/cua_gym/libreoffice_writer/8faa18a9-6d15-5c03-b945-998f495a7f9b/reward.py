"""
Reward Script: Merge cells A1-D1 in invoice table, center-align 'INVOICE' with 18pt bold font
Task ID: writer_tm_020
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Cells A1-D1 are merged (gridSpan=4)
  Component 2 (0.30): Text is 'INVOICE' and center-aligned
  Component 3 (0.35): Font is 18pt and bold
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_020'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table has at least 10 rows
    if len(table.rows) < 10:
        print(f"CRITICAL: Table has {len(table.rows)} rows, expected at least 10")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Cells A1-D1 are merged (gridSpan=4) — 0.35 points
    # In initial_env: gridSpan is None (no merge). In golden_env: gridSpan=4.
    try:
        cell_a1 = table.cell(0, 0)
        tc = cell_a1._tc
        tc_pr = tc.find(qn('w:tcPr'))
        grid_span_val = None
        if tc_pr is not None:
            gs = tc_pr.find(qn('w:gridSpan'))
            if gs is not None:
                grid_span_val = gs.get(qn('w:val'))

        if grid_span_val is not None and int(grid_span_val) >= 4:
            print(f"PASS: Component 1 — Cells A1-D1 merged with gridSpan={grid_span_val} (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — Expected gridSpan>=4, found gridSpan={grid_span_val}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Text is 'INVOICE' and center-aligned — 0.30 points
    # In initial_env: alignment=LEFT. In golden_env: alignment=CENTER.
    try:
        cell_a1 = table.cell(0, 0)
        cell_text = cell_a1.text.strip()
        # Check alignment on the paragraph inside the merged cell
        para = cell_a1.paragraphs[0]
        alignment = para.paragraph_format.alignment

        text_ok = (cell_text == 'INVOICE')
        center_ok = (alignment is not None and
                     alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

        if text_ok and center_ok:
            print(f"PASS: Component 2 — Text is 'INVOICE' and center-aligned (0.30 pts)")
            total_score += 0.30
        else:
            if not text_ok:
                print(f"FAIL: Component 2 — Expected text 'INVOICE', found '{cell_text}'")
            if not center_ok:
                print(f"FAIL: Component 2 — Expected CENTER alignment, found {alignment}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Font is 18pt and bold — 0.35 points
    # In initial_env: bold=False, size=12pt. In golden_env: bold=True, size=18pt.
    try:
        cell_a1 = table.cell(0, 0)
        runs = cell_a1.paragraphs[0].runs

        if len(runs) == 0:
            print("FAIL: Component 3 — No runs found in merged cell")
        else:
            # Check all runs in the cell for bold and size
            all_bold = all(r.font.bold is True for r in runs if r.text.strip())
            all_18pt = all(
                r.font.size is not None and abs(r.font.size.pt - 18.0) < 0.5
                for r in runs if r.text.strip()
            )

            if all_bold and all_18pt:
                sizes = [r.font.size.pt for r in runs if r.text.strip() and r.font.size]
                print(f"PASS: Component 3 — Font is bold and 18pt (sizes={sizes}) (0.35 pts)")
                total_score += 0.35
            else:
                for r in runs:
                    if r.text.strip():
                        sz = r.font.size.pt if r.font.size else None
                        print(f"  Run '{r.text}': bold={r.font.bold}, size_pt={sz}")
                if not all_bold:
                    print(f"FAIL: Component 3 — Not all runs are bold")
                if not all_18pt:
                    print(f"FAIL: Component 3 — Not all runs are 18pt")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
