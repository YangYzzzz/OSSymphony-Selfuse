"""
Reward Script: Heading 1 style page break before
Task ID: writer_tech_069
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Heading 1 style definition has page_break_before = True
  Component 2 (0.3): All Heading 1 paragraphs effectively have page_break_before (none override to False)
  Component 3 (0.2): Document integrity — still has 6 Heading 1 paragraphs
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_069'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the Heading 1 style has page_break_before enabled.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Heading 1 style definition has page_break_before = True (0.5 points)
    try:
        h1_style = None
        for style in doc.styles:
            if style.name == 'Heading 1':
                h1_style = style
                break

        if h1_style is None:
            print("FAIL: Component 1 — Heading 1 style not found in document")
        else:
            style_pbf = h1_style.paragraph_format.page_break_before
            if style_pbf is True:
                print(f"PASS: Component 1 — Heading 1 style page_break_before = {style_pbf} (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 1 — Heading 1 style page_break_before = {style_pbf}, expected True")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All Heading 1 paragraphs effectively have page_break_before (0.3 points)
    # Individual paragraphs should NOT override the style to disable page_break_before.
    # para.paragraph_format.page_break_before returns None (inherit from style) or True — both OK.
    # Only False (explicit override disabling it) would be a problem.
    try:
        h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
        if not h1_paras:
            print("FAIL: Component 2 — No Heading 1 paragraphs found")
        else:
            all_effective = True
            for para in h1_paras:
                pbf = para.paragraph_format.page_break_before
                # None means inherit from style, True means explicitly set — both are OK
                # Only False would mean an override disabling the style setting
                if pbf is False:
                    all_effective = False
                    print(f"FAIL: Component 2 — Heading 1 paragraph '{para.text[:40]}' overrides page_break_before to False")

            if all_effective:
                # Verify the style is also True (compound check with Component 1)
                style_pbf = None
                for style in doc.styles:
                    if style.name == 'Heading 1':
                        style_pbf = style.paragraph_format.page_break_before
                        break
                if style_pbf is True:
                    print(f"PASS: Component 2 — All {len(h1_paras)} Heading 1 paragraphs inherit page_break_before from style (0.3 pts)")
                    total_score += 0.3
                else:
                    print(f"FAIL: Component 2 — Style page_break_before is not True, so inheritance doesn't help")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document integrity — still has 6 Heading 1 paragraphs (0.2 points)
    # This is a compound check: we verify the style setting is True AND the document still has 6 headings.
    # On initial_env, the style is False, so this component fails there too.
    try:
        h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
        h1_count = len(h1_paras)

        # Compound check: style must have page_break_before AND 6 headings must exist
        style_pbf = None
        for style in doc.styles:
            if style.name == 'Heading 1':
                style_pbf = style.paragraph_format.page_break_before
                break

        if style_pbf is True and h1_count == 6:
            print(f"PASS: Component 3 — Document has {h1_count} Heading 1 paragraphs with style page_break_before enabled (0.2 pts)")
            total_score += 0.2
        elif style_pbf is not True:
            print(f"FAIL: Component 3 — Heading 1 style page_break_before is {style_pbf}, not True")
        else:
            print(f"FAIL: Component 3 — Expected 6 Heading 1 paragraphs, found {h1_count}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
