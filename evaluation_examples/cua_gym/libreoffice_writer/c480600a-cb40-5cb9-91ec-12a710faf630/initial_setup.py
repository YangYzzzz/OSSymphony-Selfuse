"""
Initial Setup: Set up document so Heading 1 paragraphs always start on a new page
Task ID: writer_tech_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Style configuration (ensure Heading 1 does NOT have page break before) --
    h1_style = doc.styles['Heading 1']
    h1_style.paragraph_format.page_break_before = False

    # ============================================================
    # Section 1: Introduction
    # ============================================================
    doc.add_heading('Introduction to Cloud-Native Architecture', level=1)

    doc.add_paragraph(
        'Cloud-native architecture represents a fundamental shift in how modern '
        'software systems are designed, deployed, and operated. Unlike traditional '
        'monolithic approaches, cloud-native applications are built as a collection '
        'of loosely coupled microservices that can be independently developed, '
        'tested, and scaled.'
    )
    doc.add_paragraph(
        'Organizations adopting cloud-native practices report significant improvements '
        'in deployment frequency, with some teams shipping updates multiple times per day. '
        'The Netflix engineering team, for example, deploys thousands of times daily across '
        'their microservices fleet, serving over 230 million subscribers worldwide.'
    )
    doc.add_paragraph(
        'This document provides a comprehensive overview of cloud-native principles, '
        'covering container orchestration, service mesh patterns, observability strategies, '
        'and security best practices for distributed systems.'
    )

    # ============================================================
    # Section 2: Container Orchestration
    # ============================================================
    doc.add_heading('Container Orchestration with Kubernetes', level=1)

    doc.add_paragraph(
        'Kubernetes has emerged as the de facto standard for container orchestration, '
        'managing workloads across clusters of machines. Originally developed at Google '
        'based on their internal Borg system, Kubernetes automates deployment, scaling, '
        'and management of containerized applications.'
    )

    # Add a table for Kubernetes components
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Role', 'Typical Resource Allocation']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h

    k8s_data = [
        ['API Server', 'Central management hub', '2 vCPU / 4 GB RAM'],
        ['etcd', 'Distributed key-value store', '2 vCPU / 8 GB RAM'],
        ['Scheduler', 'Pod placement decisions', '1 vCPU / 2 GB RAM'],
        ['Controller Manager', 'Reconciliation loops', '1 vCPU / 2 GB RAM'],
    ]
    for r, row_data in enumerate(k8s_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph(
        'Production clusters at Shopify handle over 80,000 pods across 6,000 nodes, '
        'processing $7.5 billion in Gross Merchandise Volume during peak events like '
        'Black Friday and Cyber Monday 2024.'
    )

    # ============================================================
    # Section 3: Service Mesh
    # ============================================================
    doc.add_heading('Service Mesh Patterns and Implementation', level=1)

    doc.add_paragraph(
        'A service mesh provides a dedicated infrastructure layer for handling '
        'service-to-service communication. By deploying sidecar proxies alongside '
        'each service instance, the mesh transparently handles traffic management, '
        'security, and observability without requiring changes to application code.'
    )
    doc.add_paragraph(
        'Istio and Linkerd are the two most widely adopted service mesh implementations. '
        'Istio, backed by Google and IBM, offers a comprehensive feature set including '
        'advanced traffic routing, mutual TLS, and fine-grained authorization policies. '
        'Linkerd, maintained by Buoyant, prioritizes simplicity and low resource overhead, '
        'with a Rust-based data plane proxy called linkerd2-proxy.'
    )
    doc.add_paragraph(
        'At Airbnb, the service mesh handles approximately 500 million internal RPC calls '
        'per minute, with P99 latency overhead below 2 milliseconds per hop. Their mesh '
        'configuration manages traffic splitting for canary deployments across 1,200 '
        'production services.'
    )

    # ============================================================
    # Section 4: Observability
    # ============================================================
    doc.add_heading('Observability in Distributed Systems', level=1)

    doc.add_paragraph(
        'Observability encompasses three pillars: metrics, logs, and traces. In distributed '
        'systems, understanding system behavior requires correlating signals across all three '
        'dimensions. Modern observability platforms like Datadog, Grafana, and Honeycomb '
        'provide unified views that help engineers quickly identify and resolve issues.'
    )

    # Add a bullet list for key metrics
    doc.add_paragraph('Key observability metrics to monitor:', style='List Bullet')
    doc.add_paragraph('Request latency percentiles (P50, P95, P99)', style='List Bullet')
    doc.add_paragraph('Error rate by service and endpoint', style='List Bullet')
    doc.add_paragraph('Resource utilization (CPU, memory, network I/O)', style='List Bullet')
    doc.add_paragraph('Queue depth and processing lag', style='List Bullet')
    doc.add_paragraph('Dependency health and circuit breaker state', style='List Bullet')

    doc.add_paragraph(
        'Stripe processes over 250 million API requests per day and relies on distributed '
        'tracing with OpenTelemetry to maintain their 99.999% uptime SLA. Each trace captures '
        'the full lifecycle of a payment request across 40+ internal services.'
    )

    # ============================================================
    # Section 5: Security
    # ============================================================
    doc.add_heading('Security Best Practices for Cloud-Native Workloads', level=1)

    doc.add_paragraph(
        'Security in cloud-native environments requires a defense-in-depth approach that '
        'addresses vulnerabilities at every layer: container images, runtime configuration, '
        'network policies, and supply chain integrity. The shared responsibility model means '
        'that while cloud providers secure the infrastructure, application teams must secure '
        'their workloads and data.'
    )
    doc.add_paragraph(
        'Zero-trust networking eliminates implicit trust within the cluster perimeter. Every '
        'request must be authenticated and authorized, regardless of its origin. Mutual TLS '
        'between services ensures both identity verification and encryption in transit. '
        'HashiCorp Vault or AWS Secrets Manager handle dynamic secret rotation, reducing the '
        'risk of credential exposure.'
    )
    doc.add_paragraph(
        'Image scanning tools like Trivy and Snyk Container analyze container images for '
        'known CVEs before deployment. Capital One reported a 73% reduction in critical '
        'vulnerabilities after integrating automated scanning into their CI/CD pipeline, '
        'catching issues an average of 14 days earlier in the development lifecycle.'
    )

    # ============================================================
    # Section 6: Future Directions
    # ============================================================
    doc.add_heading('Future Directions and Emerging Trends', level=1)

    doc.add_paragraph(
        'The cloud-native ecosystem continues to evolve rapidly. WebAssembly (Wasm) is '
        'emerging as a lightweight alternative to containers, offering near-native performance '
        'with a smaller footprint and faster startup times. Projects like Spin and Fermyon '
        'Cloud are pioneering Wasm-based serverless platforms that can cold-start in under '
        '1 millisecond, compared to several seconds for traditional containers.'
    )
    doc.add_paragraph(
        'Platform engineering is gaining traction as organizations recognize the need for '
        'internal developer platforms (IDPs) that abstract away infrastructure complexity. '
        'Backstage, originally developed at Spotify, provides a framework for building '
        'developer portals that centralize service catalogs, documentation, and CI/CD '
        'workflows. Surveys indicate that 78% of enterprise organizations plan to invest '
        'in platform engineering initiatives by 2026.'
    )
    doc.add_paragraph(
        'eBPF (extended Berkeley Packet Filter) is transforming observability and security '
        'by enabling programmable kernel-level instrumentation without modifying application '
        'code. Cilium, the leading eBPF-based networking solution, provides transparent '
        'encryption, load balancing, and network policy enforcement at kernel speed, '
        'processing over 10 million packets per second per node.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
