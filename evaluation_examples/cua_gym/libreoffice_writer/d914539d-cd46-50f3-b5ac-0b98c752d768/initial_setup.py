"""
Initial Setup: California state court brief without pleading paper formatting
Task ID: writer_legal_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard letter size with typical legal margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

    # --- Case Caption ---
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("SUPERIOR COURT OF THE STATE OF CALIFORNIA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("COUNTY OF LOS ANGELES")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()  # blank line

    # Case parties
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("ELENA MARCHETTI,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("                    Plaintiff,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("        vs.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("PACIFIC COAST PROPERTIES, LLC,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("a California limited liability company,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("                    Defendant.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # Case number
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("Case No. 24STCV08312")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # Document title
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("MEMORANDUM OF POINTS AND AUTHORITIES IN SUPPORT OF")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("PLAINTIFF'S MOTION FOR SUMMARY JUDGMENT")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # --- Section I ---
    p = doc.add_paragraph()
    run = p.add_run("I. INTRODUCTION")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    body_text = [
        "Plaintiff Elena Marchetti respectfully submits this memorandum of points and "
        "authorities in support of her motion for summary judgment against Defendant "
        "Pacific Coast Properties, LLC. This motion is based on this memorandum, the "
        "declaration of Elena Marchetti, the declaration of Robert Tanaka, the exhibits "
        "attached thereto, and such other matters as may be presented to the Court at or "
        "before the hearing on this motion.",

        "As set forth below, the undisputed material facts demonstrate that Defendant "
        "breached the commercial lease agreement dated March 15, 2023, by failing to "
        "maintain the premises in a habitable condition as required under California "
        "Civil Code Section 1941 and the express terms of the lease. Defendant's "
        "repeated failures to address documented structural deficiencies, persistent "
        "water intrusion, and non-functional HVAC systems rendered the leased premises "
        "unsuitable for Plaintiff's business operations.",

        "There is no triable issue of material fact, and Plaintiff is entitled to "
        "judgment as a matter of law on her claims for breach of contract and breach "
        "of the implied warranty of habitability.",
    ]

    for text in body_text:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Section II ---
    p = doc.add_paragraph()
    run = p.add_run("II. STATEMENT OF UNDISPUTED MATERIAL FACTS")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    facts = [
        "On March 15, 2023, Plaintiff and Defendant entered into a written commercial "
        "lease agreement for the premises located at 4521 Wilshire Boulevard, Suite 300, "
        "Los Angeles, California 90010 (the \"Premises\"). (Marchetti Decl., \u00b6 3; "
        "Ex. A.)",

        "Under Section 8.2 of the Lease, Defendant agreed to \"maintain the structural "
        "elements of the building, including the roof, exterior walls, and foundation, "
        "in good condition and repair.\" (Ex. A, \u00a7 8.2.)",

        "Beginning in June 2023, Plaintiff notified Defendant in writing of significant "
        "water intrusion through the ceiling of the Premises during periods of rain. "
        "(Marchetti Decl., \u00b6 7; Ex. B.)",

        "On July 12, 2023, licensed structural engineer Robert Tanaka, P.E., inspected "
        "the Premises and identified multiple points of failure in the roof membrane "
        "directly above Plaintiff's leased space. (Tanaka Decl., \u00b6\u00b6 4-8; Ex. C.)",

        "Despite receiving the Tanaka Report on July 20, 2023, Defendant failed to "
        "commence any repair work for a period of one hundred twelve (112) days. "
        "(Marchetti Decl., \u00b6 11.)",

        "On September 5, 2023, the HVAC system serving the Premises ceased functioning "
        "entirely. Plaintiff notified Defendant the same day. Defendant did not restore "
        "climate control until October 30, 2023. (Marchetti Decl., \u00b6\u00b6 13-14; Ex. D.)",

        "As a direct result of Defendant's failures, Plaintiff was unable to operate her "
        "interior design consulting business for a cumulative period of forty-seven (47) "
        "business days between June 2023 and November 2023. (Marchetti Decl., \u00b6 18.)",
    ]

    for i, fact in enumerate(facts, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(f"{i}. {fact}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Section III ---
    p = doc.add_paragraph()
    run = p.add_run("III. ARGUMENT")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("A. Standard of Review")
    run.bold = True
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    argument_text = [
        "A party is entitled to summary judgment when \"all the papers submitted show "
        "that there is no triable issue as to any material fact and that the moving "
        "party is entitled to a judgment as a matter of law.\" (Code Civ. Proc., "
        "\u00a7 437c, subd. (c).) The moving party bears the initial burden of making a "
        "prima facie showing that there are no triable issues of material fact. "
        "(Aguilar v. Atlantic Richfield Co. (2001) 25 Cal.4th 826, 850.)",

        "Once the moving party has met this burden, the burden shifts to the opposing "
        "party to show the existence of a triable issue of material fact. (Code Civ. "
        "Proc., \u00a7 437c, subd. (p)(2).) The opposing party may not rely on the "
        "allegations in the pleadings but must set forth specific facts showing a "
        "triable issue. (Ibid.)",
    ]

    for text in argument_text:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("B. Defendant Breached the Lease Agreement")
    run.bold = True
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    breach_text = [
        "The elements of a breach of contract claim are: (1) the existence of a "
        "contract; (2) plaintiff's performance or excuse for nonperformance; (3) "
        "defendant's breach; and (4) resulting damages. (Oasis West Realty, LLC v. "
        "Goldman (2011) 51 Cal.4th 811, 821.)",

        "Here, each element is established by undisputed evidence. The Lease is a "
        "valid, enforceable contract between the parties. (UMF \u00b6 1.) Plaintiff "
        "performed all obligations under the Lease, including timely payment of rent "
        "throughout the relevant period. (Marchetti Decl., \u00b6 5.) Defendant breached "
        "its express obligation under Section 8.2 to maintain the structural elements "
        "of the building by failing to repair the roof despite actual knowledge of the "
        "deficiency. (UMF \u00b6\u00b6 3-5.) Plaintiff suffered damages in the form of lost "
        "business income and costs of temporary relocation. (Marchetti Decl., \u00b6\u00b6 18-20.)",
    ]

    for text in breach_text:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Section IV ---
    p = doc.add_paragraph()
    run = p.add_run("IV. CONCLUSION")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(
        "For the foregoing reasons, Plaintiff Elena Marchetti respectfully requests "
        "that the Court grant her motion for summary judgment in its entirety and "
        "enter judgment in her favor against Defendant Pacific Coast Properties, LLC "
        "on all claims asserted in the Complaint."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # Signature block
    p = doc.add_paragraph()
    run = p.add_run("Dated: January 15, 2025")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    p = doc.add_paragraph()
    run = p.add_run("Respectfully submitted,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # blank line
    doc.add_paragraph()  # blank line

    p = doc.add_paragraph()
    run = p.add_run("_____________________________")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("DIANE FOSTER, ESQ. (SBN 247893)")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Foster & Nakamura LLP")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("633 West Fifth Street, Suite 2800")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Los Angeles, California 90071")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Telephone: (213) 555-0147")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Attorney for Plaintiff Elena Marchetti")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # NO line numbering - that's the task
    # NO left page border - that's the task

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
