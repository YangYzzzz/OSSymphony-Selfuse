"""
Initial Setup: Research paper and journal submission checklist
Task ID: osworld_multi_apps_reminder_doc_update_writer_009
Domain: libreoffice_writer

Creates:
- /home/user/Desktop/journal_submission_checklist.odt: checklist with 10 requirements
- /home/user/research_paper.odt: 15-page research paper with intentionally wrong formatting
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_reminder_doc_update_writer_009'
DESKTOP = '/home/user/Desktop'

# Paths
CHECKLIST_PATH = f'{DESKTOP}/journal_submission_checklist.odt'
PAPER_PATH = f'{WORKDIR}/research_paper.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_checklist():
    """Create journal_submission_checklist.odt on the Desktop."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Journal of Science — Manuscript Formatting Requirements', level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True

    intro = doc.add_paragraph(
        'All manuscripts submitted to the Journal of Science must comply with the following '
        'formatting requirements. Authors are required to ensure ALL requirements are met '
        'prior to submission.'
    )

    doc.add_paragraph()

    # Add 10 numbered requirements
    requirements = [
        ('1. Body Text Font',
         'All body text must be formatted in Times New Roman, 12pt. '
         'This applies to all paragraphs, captions, and table content.'),
        ('2. Section Headings (Heading 1)',
         'All major section headings must use the Heading 1 style, formatted as '
         '14pt bold. Sections include: Abstract, Introduction, Methods, Results, '
         'Discussion, Conclusion, and References.'),
        ('3. Line Spacing',
         'The entire manuscript, including abstract, body text, references, and captions, '
         'must use double line spacing (2.0). Single or 1.5 spacing is not acceptable.'),
        ('4. Page Margins',
         'All four page margins (top, bottom, left, right) must be set to exactly '
         '2.54 cm (1 inch). Manuscripts with margins differing from this will be '
         'returned for revision.'),
        ('5. Abstract Length',
         'The abstract must be between 150 and 250 words. Abstracts that are too short '
         'or too long will not be accepted. Authors should structure the abstract to '
         'include: Objectives, Methods, Results, and Conclusions.'),
        ('6. Figure Captions',
         'All figure captions must be placed BELOW the corresponding figure. '
         'Caption text must be 10pt italic. Captions must be numbered sequentially '
         '(e.g., Figure 1: ..., Figure 2: ...).'),
        ('7. References Format (Vancouver Style)',
         'References must follow the Vancouver citation style with sequential numbering. '
         'In-text citations use superscript numbers (e.g., Smith et al.^1). '
         'The reference list must use numbered format: [1] Author AA. Title. '
         'Journal Name. Year;Vol(Issue):Pages.'),
        ('8. Page Numbers',
         'Page numbers must appear in the top-right corner of every page, '
         'starting from page 1. Use Arabic numerals. Page numbers must be '
         'formatted to match the body text font.'),
        ('9. Running Header',
         "A running header must be included on every page. The running header "
         "must contain the first author's last name followed by the journal "
         "abbreviation (e.g., 'Anderson — J Sci'). The header must be right-aligned."),
        ('10. Table Titles',
         'Every table must have a title placed ABOVE the table. Table titles '
         'must be in bold. Titles should be numbered sequentially and descriptive '
         '(e.g., Table 1: Summary of patient demographics).'),
    ]

    for req_title, req_body in requirements:
        # Requirement heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(req_title)
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        # Requirement body
        body_para = doc.add_paragraph(req_body)
        body_para.paragraph_format.space_after = Pt(8)
        doc.add_paragraph()  # blank line between requirements

    # Footer note
    doc.add_paragraph()
    footer_note = doc.add_paragraph(
        'Manuscripts not meeting ALL of the above requirements will be returned '
        'to authors without review. For questions, contact editorial@journalofscience.org'
    )
    footer_note.runs[0].italic = True

    doc.save(CHECKLIST_PATH)
    print(f'Checklist created: {CHECKLIST_PATH}')


def create_research_paper():
    """Create research_paper.odt with intentionally wrong formatting (pre-task state)."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # Set wrong formatting: 3 cm margins (not 2.54), Calibri 11pt, 1.5 spacing
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # No running header (linked to previous / empty)
    header = section.header
    header.is_linked_to_previous = False
    # Leave header paragraph empty
    if header.paragraphs:
        header.paragraphs[0].clear()

    # No footer page numbers

    def add_body_para(text, spacing=1.5):
        """Add a body paragraph with Calibri 11pt and 1.5 spacing (wrong formatting)."""
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = spacing
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return para

    def add_section_heading(text):
        """Add a plain heading with no Heading 1 style (wrong formatting)."""
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.bold = True
        # Deliberately NOT using doc.add_heading() to avoid Heading 1 style
        return para

    # ---- Title Page ----
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run(
        'Longitudinal Assessment of Microbiome Diversity in Patients '
        'Receiving Targeted Immunotherapy: A Multi-Center Cohort Study'
    )
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(16)
    title_run.bold = True

    authors_para = doc.add_paragraph()
    authors_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors_run = authors_para.add_run(
        'James Anderson1, Priya Nair2, Tobias Müller3, Yuki Tanaka4, Sofia Reyes1'
    )
    authors_run.font.name = 'Calibri'
    authors_run.font.size = Pt(11)

    affil_para = doc.add_paragraph()
    affil_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil_run = affil_para.add_run(
        '1Department of Oncology, Stanford Medical Center, Stanford, CA, USA\n'
        '2Institute of Genomics, University of Mumbai, Mumbai, India\n'
        '3Charité – Universitätsmedizin Berlin, Berlin, Germany\n'
        '4Department of Immunology, Osaka University Hospital, Osaka, Japan'
    )
    affil_run.font.name = 'Calibri'
    affil_run.font.size = Pt(10)

    correspondence_para = doc.add_paragraph()
    corr_run = correspondence_para.add_run(
        'Corresponding author: James Anderson, james.anderson@stanford.edu'
    )
    corr_run.font.name = 'Calibri'
    corr_run.font.size = Pt(10)
    corr_run.italic = True

    doc.add_page_break()

    # ---- Abstract (400 words — too long, must be trimmed to 150-250) ----
    add_section_heading('Abstract')

    abstract_text = (
        'Background: The human gut microbiome plays a fundamental role in modulating '
        'immune responses, and its composition may profoundly influence the efficacy '
        'of immunotherapy treatments in oncology patients. Understanding how microbiome '
        'diversity changes during treatment is critical for optimizing therapeutic '
        'strategies and predicting clinical outcomes. '
        'Objectives: This study aimed to characterize longitudinal changes in gut '
        'microbiome diversity among patients undergoing targeted immunotherapy for '
        'advanced solid tumors, and to identify microbial signatures predictive of '
        'treatment response and adverse events. '
        'Methods: We conducted a prospective multi-center cohort study enrolling '
        '287 patients with advanced non-small cell lung cancer, melanoma, or '
        'colorectal carcinoma across four academic medical centers. Fecal samples '
        'were collected at baseline, weeks 4, 8, 16, and 24 post-treatment initiation. '
        '16S rRNA gene sequencing was performed on all samples, and alpha and beta '
        'diversity metrics were calculated. Multivariable regression models adjusted '
        'for age, sex, body mass index, prior antibiotic use, and tumor type were used. '
        'Results: Among 287 enrolled patients, 241 completed the 24-week follow-up. '
        'Patients who achieved partial or complete response demonstrated significantly '
        'higher baseline alpha diversity (Shannon index: 4.21 ± 0.87 vs 3.14 ± 0.92, '
        'p < 0.001). Longitudinal analysis showed a marked decrease in Firmicutes/Bacteroidetes '
        'ratio in responders at week 8 (p = 0.003). Adverse events including immune-related '
        'colitis were associated with depletion of Bifidobacterium and Lactobacillus species. '
        'Conclusions: Gut microbiome diversity at baseline and its longitudinal trajectory '
        'are strongly associated with immunotherapy outcomes in advanced cancer patients. '
        'These findings suggest that microbiome profiling should be integrated into '
        'pre-treatment assessment protocols, and that microbiome-targeted interventions '
        'may represent a viable strategy for improving immunotherapy efficacy and '
        'reducing treatment-associated adverse events. Future studies with larger cohorts '
        'and intervention designs are warranted.'
    )
    add_body_para(abstract_text)

    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.line_spacing = 1.5
    kw_run = kw_para.add_run('Keywords: ')
    kw_run.font.name = 'Calibri'
    kw_run.font.size = Pt(11)
    kw_run.bold = True
    kw_text = kw_para.add_run('gut microbiome, immunotherapy, cancer, diversity, cohort study')
    kw_text.font.name = 'Calibri'
    kw_text.font.size = Pt(11)

    doc.add_page_break()

    # ---- Introduction ----
    add_section_heading('1. Introduction')

    add_body_para(
        'The relationship between the gut microbiome and host immune function has '
        'emerged as one of the most active areas in translational oncology research. '
        'Accumulating evidence from preclinical models and retrospective clinical '
        'studies suggests that the composition and diversity of the gut microbial '
        'community significantly influences responses to immune checkpoint inhibitors, '
        'including anti-PD-1 and anti-CTLA-4 antibodies (Smith et al., 2019; '
        'Johnson and Park, 2020).'
    )

    add_body_para(
        'The gut microbiome modulates immune responses through multiple mechanisms, '
        'including regulation of T-cell differentiation, production of short-chain '
        'fatty acids that influence intestinal barrier integrity, and modulation of '
        'dendritic cell activation. Dysbiosis—an imbalance in microbial composition—'
        'has been linked to both reduced therapeutic efficacy and increased risk of '
        'immune-related adverse events (irAEs) in patients receiving immunotherapy '
        '(Garcia et al., 2021; Patel et al., 2022).'
    )

    add_body_para(
        'Despite growing interest, longitudinal prospective studies characterizing '
        'microbiome dynamics during immunotherapy treatment remain limited. Most '
        'published data derive from single-center retrospective analyses with small '
        'sample sizes and inconsistent methodology, limiting generalizability. '
        'A multi-center prospective approach with standardized collection protocols '
        'is essential to validate biomarker candidates and establish clinically '
        'actionable thresholds.'
    )

    add_body_para(
        'In this study, we present results from a multi-center cohort study designed '
        'to characterize longitudinal microbiome changes and their associations with '
        'treatment outcomes in patients receiving targeted immunotherapy for advanced '
        'solid tumors. We hypothesized that baseline microbial diversity and specific '
        'longitudinal trajectories would be associated with treatment response and '
        'the development of immune-related adverse events.'
    )

    doc.add_page_break()

    # ---- Methods ----
    add_section_heading('2. Methods')

    add_body_para(
        'Study Design and Patient Population: This was a prospective observational '
        'cohort study conducted at four academic medical centers: Stanford Medical Center '
        '(Stanford, CA, USA), University of Mumbai (Mumbai, India), Charité – '
        'Universitätsmedizin Berlin (Berlin, Germany), and Osaka University Hospital '
        '(Osaka, Japan). Adults aged 18 or older with histologically confirmed advanced '
        'non-small cell lung cancer (NSCLC), melanoma, or colorectal carcinoma initiating '
        'first-line or second-line immunotherapy were eligible for enrollment.'
    )

    add_body_para(
        'Sample Collection: Fecal samples were collected using validated self-collection '
        'kits (OMNIgene GUT, DNA Genotek) at five time points: baseline (prior to first '
        'treatment dose), and at weeks 4, 8, 16, and 24 post-treatment initiation. '
        'Samples were stored at -80°C within 4 hours of collection and processed in '
        'batches at a central laboratory (Stanford Genomics Core).'
    )

    add_body_para(
        'Microbiome Sequencing: DNA was extracted using the PowerSoil Pro Kit (Qiagen). '
        '16S rRNA gene amplicons targeting the V3-V4 hypervariable region were amplified '
        'using primers 341F/806R and sequenced on the Illumina MiSeq platform with paired-end '
        '250 bp reads. Bioinformatic processing was performed using QIIME2 (v2023.2). '
        'Amplicon sequence variants (ASVs) were identified using the DADA2 pipeline, '
        'and taxonomic classification was performed using the SILVA 138 reference database.'
    )

    add_body_para(
        'Statistical Analysis: Alpha diversity metrics (Shannon index, observed species, '
        'Faith\'s phylogenetic diversity) were calculated after rarefaction to 10,000 reads '
        'per sample. Beta diversity was assessed using Bray-Curtis dissimilarity and '
        'weighted UniFrac distances. Associations between microbiome diversity and '
        'treatment response were evaluated using multivariable linear mixed-effects models '
        'adjusted for age, sex, BMI, tumor type, prior antibiotic use, and center. '
        'Statistical significance threshold was set at p < 0.05.'
    )

    doc.add_page_break()

    # ---- Results ----
    add_section_heading('3. Results')

    add_body_para(
        'Patient Characteristics: A total of 287 patients were enrolled between '
        'January 2022 and June 2023. Of these, 241 (84%) completed the 24-week '
        'follow-up. Demographics and baseline characteristics are summarized in Table 1.'
    )

    # Table 1 (no title above - wrong formatting)
    # Note: in initial state, tables have no title above them
    table1_para = doc.add_paragraph()
    table1_run = table1_para.add_run('Table 1')
    table1_run.font.name = 'Calibri'
    table1_run.font.size = Pt(11)
    table1_run.bold = True
    # No caption above - just table immediately follows

    table1 = doc.add_table(rows=6, cols=4)
    table1.style = 'Table Grid'
    headers_row = ['Characteristic', 'All Patients (n=287)', 'Responders (n=143)', 'Non-Responders (n=144)']
    for j, h in enumerate(headers_row):
        cell = table1.cell(0, j)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Calibri'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    data_rows = [
        ['Age (years), mean ± SD', '58.4 ± 11.2', '57.9 ± 10.8', '58.9 ± 11.6'],
        ['Female sex, n (%)', '138 (48.1)', '72 (50.3)', '66 (45.8)'],
        ['Tumor type: NSCLC, n (%)', '124 (43.2)', '61 (42.7)', '63 (43.8)'],
        ['Tumor type: Melanoma, n (%)', '92 (32.1)', '48 (33.6)', '44 (30.6)'],
        ['Tumor type: CRC, n (%)', '71 (24.7)', '34 (23.8)', '37 (25.7)'],
    ]
    for i, row_data in enumerate(data_rows, 1):
        for j, val in enumerate(row_data):
            cell = table1.cell(i, j)
            cell.text = val
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    add_body_para(
        'Baseline Microbiome Diversity: Responders demonstrated significantly higher '
        'baseline alpha diversity compared to non-responders (Shannon index: 4.21 ± 0.87 '
        'vs 3.14 ± 0.92, p < 0.001; Figure 1). Faith\'s phylogenetic diversity was also '
        'significantly higher in responders (p = 0.004).'
    )

    # Figure 1 caption on the side (wrong position and formatting)
    fig1_caption = doc.add_paragraph()
    fig1_caption.paragraph_format.line_spacing = 1.5
    fig1_cap_run = fig1_caption.add_run(
        '[Figure 1: Baseline Shannon diversity index stratified by treatment response. '
        'Error bars represent standard deviation.]'
    )
    fig1_cap_run.font.name = 'Calibri'
    fig1_cap_run.font.size = Pt(11)
    # Note: NOT italic, NOT 10pt - wrong formatting for figure caption

    add_body_para(
        'Longitudinal Dynamics: Longitudinal analysis revealed distinct microbiome '
        'trajectories between responders and non-responders. At week 8, responders '
        'demonstrated a significant decrease in the Firmicutes/Bacteroidetes ratio '
        '(from 2.31 at baseline to 1.48 at week 8, p = 0.003), while non-responders '
        'showed no significant change (from 2.28 to 2.19, p = 0.74; Figure 2).'
    )

    fig2_caption = doc.add_paragraph()
    fig2_caption.paragraph_format.line_spacing = 1.5
    fig2_cap_run = fig2_caption.add_run(
        '[Figure 2: Longitudinal changes in Firmicutes/Bacteroidetes ratio by treatment response group.]'
    )
    fig2_cap_run.font.name = 'Calibri'
    fig2_cap_run.font.size = Pt(11)
    # Note: NOT italic, NOT 10pt

    add_body_para(
        'Adverse Events: Patients who developed immune-related colitis (n = 34, 11.8%) '
        'showed significantly depleted Bifidobacterium (mean relative abundance: 0.8% vs '
        '3.2%, p < 0.001) and Lactobacillus (0.4% vs 1.8%, p = 0.002) species at baseline '
        'compared to patients without colitis.'
    )

    doc.add_page_break()

    # ---- Discussion ----
    add_section_heading('4. Discussion')

    add_body_para(
        'This multi-center prospective cohort study provides robust longitudinal evidence '
        'that gut microbiome diversity is significantly associated with immunotherapy '
        'outcomes in patients with advanced solid tumors. Our finding that higher baseline '
        'alpha diversity predicts treatment response is consistent with prior retrospective '
        'analyses (Gopalakrishnan et al., 2018; Routy et al., 2018) and extends this '
        'observation to a diverse international patient population with standardized '
        'collection protocols.'
    )

    add_body_para(
        'The significant decrease in Firmicutes/Bacteroidetes ratio observed exclusively '
        'in responders at week 8 suggests a treatment-microbiome interaction that may '
        'reflect immune activation of the mucosal immune system. This is consistent with '
        'preclinical data showing that immunotherapy-mediated tumor rejection is associated '
        'with enrichment of specific short-chain fatty acid-producing bacteria (Smith et al., '
        '2019). The clinical implications of this finding require validation in interventional '
        'studies using pre- and probiotic approaches.'
    )

    add_body_para(
        'The association between depleted Bifidobacterium and Lactobacillus at baseline and '
        'subsequent development of immune-related colitis is a novel finding with important '
        'clinical implications. These commensal organisms play key roles in maintaining '
        'intestinal barrier integrity and modulating mucosal immune responses. Their depletion '
        'may represent a pre-treatment vulnerability that could be addressed through targeted '
        'microbiome modulation strategies.'
    )

    add_body_para(
        'Several limitations of this study should be acknowledged. First, as an observational '
        'study, we cannot establish causality between microbiome changes and treatment outcomes. '
        'Second, dietary information was not systematically collected, which may confound '
        'microbiome-outcome associations. Third, the 16S rRNA approach provides limited '
        'taxonomic resolution compared to shotgun metagenomic sequencing. Future studies '
        'should address these limitations.'
    )

    doc.add_page_break()

    # ---- Conclusion ----
    add_section_heading('5. Conclusion')

    add_body_para(
        'This prospective multi-center cohort study demonstrates that gut microbiome '
        'diversity at baseline and its longitudinal trajectory during immunotherapy '
        'treatment are significantly associated with treatment response and adverse '
        'events. These findings support the integration of microbiome profiling into '
        'clinical decision-making and the development of microbiome-targeted interventions '
        'to improve immunotherapy outcomes in advanced cancer patients.'
    )

    doc.add_page_break()

    # ---- Tables (additional) ----
    add_section_heading('6. Tables')

    # Table 2 label (wrong: should be above in bold, but here it's inline)
    t2_note = doc.add_paragraph()
    t2_run = t2_note.add_run('Table 2')
    t2_run.font.name = 'Calibri'
    t2_run.font.size = Pt(11)
    t2_run.bold = True

    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    t2_headers = ['Microbial Genus', 'Responders Mean (%)', 'Non-Responders Mean (%)']
    for j, h in enumerate(t2_headers):
        cell = table2.cell(0, j)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Calibri'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    t2_data = [
        ['Bifidobacterium', '3.2 ± 1.1', '1.4 ± 0.8'],
        ['Lactobacillus', '2.8 ± 0.9', '1.1 ± 0.6'],
        ['Akkermansia', '4.1 ± 1.8', '2.3 ± 1.3'],
        ['Faecalibacterium', '6.7 ± 2.2', '4.9 ± 2.0'],
    ]
    for i, row_data in enumerate(t2_data, 1):
        for j, val in enumerate(row_data):
            cell = table2.cell(i, j)
            cell.text = val
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    doc.add_page_break()

    # ---- References (APA style — wrong, should be Vancouver) ----
    add_section_heading('References')

    references_apa = [
        'Garcia, M. L., Thompson, R. K., & Williams, S. D. (2021). Gut dysbiosis and '
        'immune-related adverse events in checkpoint inhibitor therapy. Nature Immunology, '
        '22(4), 412-425.',
        'Gopalakrishnan, V., Spencer, C. N., Nezi, L., & Wargo, J. A. (2018). Gut '
        'microbiome modulates response to anti-PD-1 immunotherapy in melanoma patients. '
        'Science, 359(6371), 97-103.',
        'Johnson, P. R., & Park, S. Y. (2020). The role of Bacteroides in anti-CTLA-4 '
        'therapy efficacy. Cell Host & Microbe, 27(3), 345-358.',
        'Patel, R. S., Kumar, A., & Singh, N. (2022). Microbiome modulation strategies '
        'to improve immunotherapy outcomes: A systematic review. Lancet Oncology, '
        '23(8), 1024-1039.',
        'Routy, B., Le Chatelier, E., Derosa, L., & Zitvogel, L. (2018). Gut microbiome '
        'influences efficacy of PD-1 based immunotherapy against epithelial tumors. '
        'Science, 359(6371), 91-97.',
        'Smith, J. A., Brown, K. L., & Davis, H. (2019). Short-chain fatty acids and '
        'tumor immunity: Links between gut microbiota and anti-tumor immune responses. '
        'Immunity, 51(2), 283-297.',
    ]

    for i, ref in enumerate(references_apa, 1):
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.line_spacing = 1.5
        ref_para.paragraph_format.left_indent = Cm(1.0)
        ref_para.paragraph_format.first_line_indent = Cm(-1.0)
        ref_run = ref_para.add_run(ref)
        ref_run.font.name = 'Calibri'
        ref_run.font.size = Pt(11)

    doc.save(PAPER_PATH)
    print(f'Research paper created: {PAPER_PATH}')


def main():
    create_checklist()
    create_research_paper()

    # GUI-ready startup: open checklist first (the task says to open it from Desktop)
    # Then open the paper that needs to be edited
    launch_gui(f'libreoffice --writer "{CHECKLIST_PATH}"', delay_sec=2.5)
    launch_gui(f'libreoffice --writer "{PAPER_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with both documents on DISPLAY=:0')


main()
