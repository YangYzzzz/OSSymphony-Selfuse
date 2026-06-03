"""
Reward Script: Apply journal formatting requirements to research_paper.odt
Task ID: osworld_multi_apps_reminder_doc_update_writer_009
Domain: libreoffice_writer

Scoring Rubric (total 1.0):
  Component 1: Body text font changed to Times New Roman 12pt              (0.15)
  Component 2: Line spacing changed to double (2.0)                       (0.15)
  Component 3: Margins changed to 2.54cm on all sides                     (0.10)
  Component 4: Heading 1 style applied to section titles (14pt bold TNR)  (0.15)
  Component 5: Figure captions are 10pt italic below figures              (0.10)
  Component 6: References in Vancouver style (numbered [1], [2], ...)     (0.15)
  Component 7: Running header with first author's last name + page number (0.20)

Total: 1.0
"""

import os
import re
import zipfile

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_reminder_doc_update_writer_009'

# Tolerance for margin comparison (EMU): 36000 EMU = 0.04cm
MARGIN_TOLERANCE_EMU = 50000  # ~0.13cm tolerance


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Pt, Emu
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Body text font changed to Times New Roman 12pt (0.15 points)
    # Initial state: Calibri 11pt (139700 EMU). Golden state: Times New Roman 12pt (152400 EMU).
    try:
        tnr_12pt_body_count = 0
        body_run_count = 0
        # Check body paragraphs (Normal style) for font name and size
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            # Body text paragraphs (Normal style, not headings)
            if style_name in ('Normal', 'Default Paragraph Style') or not style_name:
                for run in para.runs:
                    if run.text.strip() and len(run.text.strip()) > 3:
                        body_run_count += 1
                        font_name = run.font.name
                        font_size = run.font.size
                        name_ok = font_name and 'Times New Roman' in font_name
                        # 12pt = 152400 EMU
                        size_ok = font_size and abs(font_size - Pt(12)) < 10000
                        if name_ok and size_ok:
                            tnr_12pt_body_count += 1

        if body_run_count > 0:
            ratio = tnr_12pt_body_count / body_run_count
            if ratio >= 0.7:
                print(f"PASS: Component 1 — Body text is Times New Roman 12pt ({tnr_12pt_body_count}/{body_run_count} runs, {ratio:.1%}) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 — Expected Times New Roman 12pt body text, only {tnr_12pt_body_count}/{body_run_count} runs match ({ratio:.1%})")
        else:
            print("FAIL: Component 1 — No body run text found for font check")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Line spacing changed to double (2.0) (0.15 points)
    # Initial state: 1.5 spacing. Golden state: 2.0 (double) spacing.
    try:
        double_spaced_count = 0
        total_content_paras = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            total_content_paras += 1
            pf = para.paragraph_format
            ls = pf.line_spacing
            # line_spacing == 2.0 means double spacing (float) or could be Pt value
            # Double spacing float value is 2.0
            if ls is not None:
                if isinstance(ls, (int, float)) and abs(ls - 2.0) < 0.1:
                    double_spaced_count += 1
                elif hasattr(ls, 'pt') and abs(ls.pt - 2.0) < 0.1:
                    double_spaced_count += 1

        if total_content_paras > 0:
            ratio = double_spaced_count / total_content_paras
            if ratio >= 0.6:
                print(f"PASS: Component 2 — Double line spacing applied ({double_spaced_count}/{total_content_paras} content paragraphs, {ratio:.1%}) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — Expected double spacing (2.0), only {double_spaced_count}/{total_content_paras} paragraphs have it ({ratio:.1%})")
        else:
            print("FAIL: Component 2 — No content paragraphs found for spacing check")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Margins changed to 2.54cm on all sides (0.10 points)
    # Initial state: 3.0cm. Golden state: 2.54cm (1 inch).
    # 2.54cm = 1 inch = 914400 EMU
    try:
        target_margin_emu = 914400  # 2.54cm in EMU
        section = doc.sections[0]
        lm = section.left_margin
        rm = section.right_margin
        tm = section.top_margin
        bm = section.bottom_margin

        if all(m is not None for m in [lm, rm, tm, bm]):
            margins_ok = (
                abs(lm - target_margin_emu) <= MARGIN_TOLERANCE_EMU and
                abs(rm - target_margin_emu) <= MARGIN_TOLERANCE_EMU and
                abs(tm - target_margin_emu) <= MARGIN_TOLERANCE_EMU and
                abs(bm - target_margin_emu) <= MARGIN_TOLERANCE_EMU
            )
            if margins_ok:
                print(f"PASS: Component 3 — All margins set to ~2.54cm (left={lm/914400*2.54:.2f}cm, right={rm/914400*2.54:.2f}cm, top={tm/914400*2.54:.2f}cm, bottom={bm/914400*2.54:.2f}cm) (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 3 — Expected 2.54cm margins. Got: left={lm/914400*2.54:.2f}cm, right={rm/914400*2.54:.2f}cm, top={tm/914400*2.54:.2f}cm, bottom={bm/914400*2.54:.2f}cm")
        else:
            print("FAIL: Component 3 — Could not read margin values from sections")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Heading 1 style applied to section titles, 14pt bold (0.15 points)
    # Initial state: Normal style, no headings applied.
    # Golden state: 'Heading 1' style applied with 14pt bold Times New Roman.
    try:
        heading1_paras = [p for p in doc.paragraphs if p.style and 'Heading 1' in p.style.name]
        # Expected headings: Abstract, 1. Introduction, 2. Methods, 3. Results, 4. Discussion, 5. Conclusion, 6. Tables, References
        if len(heading1_paras) >= 4:
            # Verify font properties of heading 1 style or inline
            h1_style = None
            try:
                h1_style = doc.styles['Heading 1']
            except Exception:
                pass

            h1_font_ok = False
            # Check style-level font
            if h1_style and h1_style.font.size and abs(h1_style.font.size.pt - 14.0) < 0.5 and h1_style.font.bold:
                h1_font_ok = True
            else:
                # Check first heading paragraph runs
                if heading1_paras:
                    first_h1 = heading1_paras[0]
                    for r in first_h1.runs:
                        if r.text.strip():
                            sz = r.font.size
                            b = r.font.bold
                            if sz and abs(sz.pt - 14.0) < 0.5 and b:
                                h1_font_ok = True
                                break

            if h1_font_ok:
                print(f"PASS: Component 4 — Heading 1 style applied to {len(heading1_paras)} section titles with 14pt bold font (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — Heading 1 style applied ({len(heading1_paras)} paras) but font properties not meeting 14pt bold requirement")
        else:
            print(f"FAIL: Component 4 — Expected ≥4 paragraphs with Heading 1 style, found {len(heading1_paras)}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Figure captions are 10pt italic below figures (0.10 points)
    # Initial state: captions inline with figures, not italic.
    # Golden state: dedicated caption paragraphs with 10pt italic font (127000 EMU = 10pt).
    try:
        # Look for paragraphs containing "Figure" that are italic and 10pt
        figure_caption_paras = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith('Figure ') or text.startswith('[Figure '):
                # Check if it's a standalone caption (not a placeholder like [Figure 1])
                # Captions contain descriptive text after ':'
                if ':' in text and len(text) > 20:
                    # Check runs for italic 10pt
                    for r in para.runs:
                        if r.text.strip():
                            sz = r.font.size
                            it = r.font.italic
                            # 10pt = 127000 EMU
                            size_ok = sz and abs(sz.pt - 10.0) < 0.5
                            if it and size_ok:
                                figure_caption_paras.append(para)
                                break

        if len(figure_caption_paras) >= 1:
            print(f"PASS: Component 5 — Found {len(figure_caption_paras)} figure caption(s) with 10pt italic formatting (0.10 pts)")
            total_score += 0.10
        else:
            # Maybe check without size strictness
            italic_figure_captions = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if (text.startswith('Figure ') or text.startswith('[Figure ')) and ':' in text and len(text) > 20:
                    for r in para.runs:
                        if r.text.strip() and r.font.italic:
                            italic_figure_captions.append(para)
                            break
            if italic_figure_captions:
                print(f"FAIL: Component 5 — Found {len(italic_figure_captions)} italic figure caption(s) but size is not 10pt")
            else:
                print(f"FAIL: Component 5 — No figure captions found with 10pt italic formatting")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: References in Vancouver style - numbered [1], [2], ... (0.15 points)
    # Initial state: APA style (Author, A. B., Year. Title...).
    # Golden state: Vancouver style ([1] Author AB. Title...).
    try:
        vancouver_refs = []
        in_references = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Detect References section
            if text.lower() in ('references', '6. references', 'references:'):
                in_references = True
                continue
            if in_references:
                # Vancouver format: starts with [N] where N is a number
                if re.match(r'^\[\d+\]', text):
                    vancouver_refs.append(text)

        if len(vancouver_refs) >= 3:
            print(f"PASS: Component 6 — Found {len(vancouver_refs)} Vancouver-style numbered references ([1], [2], ...) (0.15 pts)")
            total_score += 0.15
        else:
            # Check if references exist but in wrong format
            apa_refs = []
            in_references2 = False
            for para in doc.paragraphs:
                text = para.text.strip()
                if text.lower() in ('references', '6. references', 'references:'):
                    in_references2 = True
                    continue
                if in_references2 and text:
                    # APA style: Author, A. B. (Year). Title.
                    if re.match(r'^[A-Z][a-z]+,\s+[A-Z]', text):
                        apa_refs.append(text)
            if apa_refs:
                print(f"FAIL: Component 6 — References are in APA format ({len(apa_refs)} refs found), not Vancouver [N] style")
            else:
                print(f"FAIL: Component 6 — Only {len(vancouver_refs)} Vancouver-style references found (need ≥3)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Running header with first author's last name + page number (top-right) (0.20 points)
    # Initial state: Empty header.
    # Golden state: Header contains "Anderson" (first author last name) and a PAGE field code.
    try:
        section = doc.sections[0]
        header = section.header
        header_has_author = False
        header_has_page_num = False
        header_text = ''

        if header and header.paragraphs:
            for para in header.paragraphs:
                header_text += para.text

            # Check for author last name (Anderson) in header text
            if 'Anderson' in header_text:
                header_has_author = True

            # Check for PAGE field code in header XML (page number)
            # The header XML contains <w:instrText> PAGE </w:instrText>
            try:
                with zipfile.ZipFile(file_path, 'r') as z:
                    if 'word/header1.xml' in z.namelist():
                        header_xml = z.read('word/header1.xml').decode('utf-8')
                        if 'PAGE' in header_xml and ('fldChar' in header_xml or 'instrText' in header_xml):
                            header_has_page_num = True
            except Exception as xml_e:
                print(f"  NOTE: Could not read header XML: {xml_e}")
                # Fallback: check if header text has digits (cached page number)
                if any(c.isdigit() for c in header_text):
                    header_has_page_num = True

        if header_has_author and header_has_page_num:
            print(f"PASS: Component 7 — Header contains author name ('Anderson') and page number field. Header text: {repr(header_text[:60])} (0.20 pts)")
            total_score += 0.20
        elif header_has_author:
            print(f"FAIL: Component 7 — Header has author name but missing page number field. Header: {repr(header_text[:60])}")
        elif header_has_page_num:
            print(f"FAIL: Component 7 — Header has page number but missing first author's last name. Header: {repr(header_text[:60])}")
        else:
            print(f"FAIL: Component 7 — Header is missing both author name and page number. Header text: {repr(header_text[:60])}")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Run verification against the canonical artifact path
file_path = f'{WORKDIR}/research_paper.odt'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
