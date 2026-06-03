"""
Reward Script: Remove all tab stops from paragraphs 3-7 and set a single decimal-aligned
               tab stop at 12 cm for the financial figures.
Task ID: writer_para_039
Domain: libreoffice_writer
Scoring:
  Component 1 (0.70): Paragraphs 3-7 each have exactly one tab stop, DECIMAL-aligned,
                       positioned at approximately 12 cm (within 0.2 cm tolerance).
  Component 2 (0.30): Paragraphs 3-7 have NO remaining LEFT-aligned tab stop
                       (the original 8 cm LEFT tab stops are fully removed).
  Total: 1.0
"""

import os
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Cm

WORKDIR = '/home/user'
TASK_ID = 'writer_para_039'

# 12 cm in EMU (1 inch = 914400 EMU; 1 cm = 914400/2.54 EMU)
CM_PER_EMU = 2.54 / 914400
TARGET_CM = 12.0
TOLERANCE_CM = 0.2   # allow ±0.2 cm rounding from docx/LibreOffice

# Paragraph indices (0-based) that must be modified: paragraphs 3–7 → indices 2–6
MODIFIED_PARA_INDICES = [2, 3, 4, 5, 6]
# Expected financial paragraph texts (pre-condition sanity check)
EXPECTED_TEXTS = [
    'Personnel Costs',
    'Equipment & Supplies',
    'Travel & Conferences',
    'Software Licenses',
    'Miscellaneous',
]


def get_meaningful_tab_stops(para):
    """Return tab stops, filtering out CLEAR and default LEFT@0 entries."""
    result = []
    for ts in para.paragraph_format.tab_stops:
        if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
            continue
        if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
            continue
        result.append(ts)
    return result


def verify_task(file_path):
    """
    Verify that paragraphs 3-7 have their old tab stops replaced with a single
    DECIMAL-aligned tab stop at 12 cm.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Pre-condition: verify we have at least 8 paragraphs
    if len(doc.paragraphs) < 8:
        print(f"CRITICAL: Expected >= 8 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Pre-condition: verify financial paragraphs contain expected text
    for i, expected_prefix in zip(MODIFIED_PARA_INDICES, EXPECTED_TEXTS):
        actual_text = doc.paragraphs[i].text
        if not actual_text.startswith(expected_prefix):
            print(f"CRITICAL: Para {i+1} text mismatch — expected to start with '{expected_prefix}', got '{actual_text[:50]}'")
            print("REWARD: 0.0")
            return 0.0

    # Component 1: Each of paragraphs 3-7 must have EXACTLY ONE tab stop,
    #              DECIMAL-aligned, positioned at ~12 cm (0.70 points)
    #
    # This check FAILS on initial_env (tab stops are LEFT@8cm, not DECIMAL@12cm)
    # and PASSES on golden_env (tab stops are DECIMAL@12cm).
    try:
        comp1_passed_count = 0
        comp1_total = len(MODIFIED_PARA_INDICES)

        for idx in MODIFIED_PARA_INDICES:
            para = doc.paragraphs[idx]
            tabs = get_meaningful_tab_stops(para)

            has_exactly_one = len(tabs) == 1
            if not has_exactly_one:
                print(f"FAIL: Component 1 — Para {idx+1} has {len(tabs)} tab stop(s), expected 1; tabs={[(str(t.alignment), t.position) for t in tabs]}")
                continue

            ts = tabs[0]
            is_decimal = (ts.alignment == WD_TAB_ALIGNMENT.DECIMAL)
            pos_cm = ts.position * CM_PER_EMU
            is_at_12cm = abs(pos_cm - TARGET_CM) <= TOLERANCE_CM

            if is_decimal and is_at_12cm:
                print(f"PASS: Component 1 — Para {idx+1} has DECIMAL tab at {pos_cm:.2f} cm (EMU={ts.position})")
                comp1_passed_count += 1
            else:
                if not is_decimal:
                    print(f"FAIL: Component 1 — Para {idx+1} tab alignment is {ts.alignment}, expected DECIMAL")
                if not is_at_12cm:
                    print(f"FAIL: Component 1 — Para {idx+1} tab position is {pos_cm:.2f} cm, expected ~{TARGET_CM} cm")

        if comp1_passed_count == comp1_total:
            print(f"PASS: Component 1 — All {comp1_total} paragraphs have DECIMAL tab at ~12 cm (0.70 pts)")
            total_score += 0.70
        else:
            print(f"FAIL: Component 1 — Only {comp1_passed_count}/{comp1_total} paragraphs correct (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraphs 3-7 must have NO left-aligned tab stop remaining.
    #              The original LEFT@8cm tab stop must be fully removed. (0.30 points)
    #
    # This check FAILS on initial_env (LEFT@8cm is present) and PASSES on golden_env
    # (LEFT tab stops are gone, replaced by DECIMAL).
    try:
        comp2_passed_count = 0
        comp2_total = len(MODIFIED_PARA_INDICES)

        for idx in MODIFIED_PARA_INDICES:
            para = doc.paragraphs[idx]
            tabs = get_meaningful_tab_stops(para)

            has_left_tab = any(
                ts.alignment == WD_TAB_ALIGNMENT.LEFT
                for ts in tabs
            )

            if not has_left_tab:
                print(f"PASS: Component 2 — Para {idx+1} has no LEFT-aligned tab stop")
                comp2_passed_count += 1
            else:
                left_tabs = [(ts.alignment, ts.position, ts.position * CM_PER_EMU) for ts in tabs if ts.alignment == WD_TAB_ALIGNMENT.LEFT]
                print(f"FAIL: Component 2 — Para {idx+1} still has LEFT tab(s): {left_tabs}")

        if comp2_passed_count == comp2_total:
            print(f"PASS: Component 2 — All {comp2_total} paragraphs have no residual LEFT tab stops (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — {comp2_total - comp2_passed_count}/{comp2_total} paragraphs still have LEFT tab stops (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Entrypoint: test the canonical artifact path
file_path = f'{WORKDIR}/budget_summary.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
