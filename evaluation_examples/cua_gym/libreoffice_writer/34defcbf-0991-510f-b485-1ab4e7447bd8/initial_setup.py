"""
Initial Setup: Remove all tab stops from paragraphs 3-7 and set decimal-aligned tab at 12cm
Task ID: writer_para_039
Domain: libreoffice_writer

Creates budget_summary.docx with:
- 8 paragraphs as per context
- Paragraphs 3-7 have various (non-decimal) tab stops that need to be replaced
- Paragraphs 1, 2, 8 have no custom tab stops
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

WORKDIR = '/home/user'
TASK_ID = 'writer_para_039'
OUTPUT = f'{WORKDIR}/budget_summary.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1 — 'Department Budget Summary FY2025'
    heading = doc.add_heading('Department Budget Summary FY2025', level=1)
    # No custom tab stops on heading

    # Paragraph 2: 'All figures in USD thousands'
    para2 = doc.add_paragraph('All figures in USD thousands')
    # No custom tab stops

    # Paragraphs 3-7: Financial line items with existing tab stops (LEFT at 8cm)
    # These need to be replaced by task with a single DECIMAL at 12cm
    financial_lines = [
        'Personnel Costs\t2,450.75',
        'Equipment & Supplies\t387.50',
        'Travel & Conferences\t125.00',
        'Software Licenses\t89.25',
        'Miscellaneous\t43.50',
    ]

    for line in financial_lines:
        para = doc.add_paragraph(line)
        # Add a LEFT-aligned tab stop at 8cm (not decimal, not at 12cm — to be replaced)
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.LEFT)

    # Paragraph 8: Note
    para8 = doc.add_paragraph(
        'Note: Budget allocations are subject to quarterly review and adjustment.'
    )
    # No custom tab stops

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
