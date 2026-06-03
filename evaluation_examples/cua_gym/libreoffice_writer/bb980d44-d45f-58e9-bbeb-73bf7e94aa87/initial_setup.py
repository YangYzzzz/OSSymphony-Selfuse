"""
Initial Setup: Create a Writer document with a Results section where a table needs to be inserted.
Task ID: writer_acad_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading('Psychometric Properties of the Workplace Engagement Scale (WES-30)', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract-like intro paragraph
    doc.add_heading('Abstract', level=2)
    p = doc.add_paragraph(
        'This study examined the psychometric properties of the Workplace Engagement Scale (WES-30), '
        'a 30-item self-report instrument designed to measure employee engagement across five dimensions: '
        'cognitive absorption, emotional commitment, behavioral effort, social connection, and professional growth. '
        'A sample of 412 full-time employees from diverse industries completed the WES-30 along with established '
        'measures of job satisfaction, organizational commitment, and burnout. Exploratory factor analysis supported '
        'a five-factor structure accounting for 67.3% of total variance. Internal consistency was strong '
        '(Cronbach\'s alpha ranging from .82 to .91 across subscales).'
    )

    # Method section
    doc.add_heading('Method', level=2)
    doc.add_heading('Participants', level=3)
    doc.add_paragraph(
        'Participants were 412 full-time employees (58% female, 42% male) recruited from organizations '
        'in the technology (n = 134), healthcare (n = 98), education (n = 87), and financial services (n = 93) '
        'sectors. Mean age was 36.4 years (SD = 9.2), and average tenure was 5.7 years (SD = 4.1).'
    )

    doc.add_heading('Measures', level=3)
    doc.add_paragraph(
        'The Workplace Engagement Scale (WES-30) consists of 30 items rated on a 7-point Likert scale '
        '(1 = strongly disagree to 7 = strongly agree). Items were developed through a combination of '
        'literature review, expert panel evaluation, and cognitive interviewing with a pilot sample of 45 employees.'
    )

    doc.add_heading('Procedure', level=3)
    doc.add_paragraph(
        'Data were collected via an online survey platform over a four-week period. Participants provided '
        'informed consent and completed the battery of measures in a single session lasting approximately '
        '25 minutes. Compensation consisted of a $15 gift card.'
    )

    # Results section - this is where the agent needs to insert the table
    doc.add_heading('Results', level=2)
    doc.add_heading('Factor Structure and Item Analysis', level=3)
    doc.add_paragraph(
        'Exploratory factor analysis using principal axis factoring with oblimin rotation was conducted on '
        'the 30 WES items. The Kaiser-Meyer-Olkin measure of sampling adequacy was .91, and Bartlett\'s test '
        'of sphericity was significant (p < .001), confirming the suitability of the data for factor analysis. '
        'Five factors with eigenvalues greater than 1.0 were extracted, accounting for 67.3% of total variance. '
        'Item-level descriptive statistics and factor loadings from the pattern matrix are presented in the table below.'
    )

    # Discussion section (after where the table should go)
    doc.add_heading('Discussion', level=2)
    doc.add_paragraph(
        'The present study provides initial psychometric evidence for the Workplace Engagement Scale (WES-30). '
        'The five-factor structure aligns with the theoretical framework proposed by Kahn (1990) and extends it '
        'by incorporating social and growth dimensions. All items demonstrated acceptable factor loadings (> .40) '
        'and contributed meaningfully to their respective subscales.'
    )

    doc.add_paragraph(
        'Convergent validity was supported by moderate-to-strong correlations with established measures of '
        'job satisfaction (r = .62, p < .001) and organizational commitment (r = .58, p < .001). Discriminant '
        'validity was evidenced by the expected negative correlation with burnout (r = -.54, p < .001). '
        'These findings suggest the WES-30 captures a construct related to but distinct from existing measures '
        'of workplace attitudes.'
    )

    doc.add_heading('References', level=2)
    doc.add_paragraph(
        'Kahn, W. A. (1990). Psychological conditions of personal engagement and disengagement at work. '
        'Academy of Management Journal, 33(4), 692-724.'
    )
    doc.add_paragraph(
        'Schaufeli, W. B., & Bakker, A. B. (2004). Job demands, job resources, and their relationship with '
        'burnout and engagement: A multi-sample study. Journal of Organizational Behavior, 25(3), 293-315.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
