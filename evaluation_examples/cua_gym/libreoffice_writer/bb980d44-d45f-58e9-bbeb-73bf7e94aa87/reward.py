"""
Reward Script: Verify survey results table in Writer document
Task ID: writer_acad_069
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Table exists with correct dimensions (31 rows x 5 cols)
  Component 2 (0.30): Correct column headers
  Component 3 (0.25): Header row repeat property enabled
  Component 4 (0.15): Data rows populated with content
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_069'

EXPECTED_HEADERS = ['Item Number', 'Question Text', 'Mean', 'SD', 'Factor Loading']


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions (0.30 points)
    # Initial env has 0 tables; golden has 1 table with 31 rows x 5 cols
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 1 — No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 31 and num_cols == 5:
                print(f"PASS: Component 1 — Table has {num_rows} rows x {num_cols} cols (0.30 pts)")
                total_score += 0.30
            else:
                # Partial credit: table exists but wrong dimensions
                if num_rows >= 2 and num_cols == 5:
                    partial = 0.15
                    print(f"PARTIAL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected 31x5 ({partial} pts)")
                    total_score += partial
                elif num_rows >= 2 and num_cols >= 3:
                    partial = 0.10
                    print(f"PARTIAL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected 31x5 ({partial} pts)")
                    total_score += partial
                else:
                    print(f"FAIL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected 31x5")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Precondition gate: need at least one table to continue
    if len(doc.tables) == 0:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    table = doc.tables[0]

    # Component 2: Correct column headers (0.30 points)
    # Initial env has no table, so no headers. Golden has specific headers.
    try:
        actual_headers = [cell.text.strip() for cell in table.rows[0].cells]
        matching = sum(1 for a, e in zip(actual_headers, EXPECTED_HEADERS) if a == e)
        if actual_headers == EXPECTED_HEADERS:
            print(f"PASS: Component 2 — All 5 column headers match (0.30 pts)")
            total_score += 0.30
        elif matching >= 3:
            partial = 0.15
            print(f"PARTIAL: Component 2 — {matching}/5 headers match ({partial} pts). Got: {actual_headers}")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Headers mismatch. Expected: {EXPECTED_HEADERS}, Got: {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row repeat property (0.25 points)
    # Initial env has no table. Golden has tblHeader on first row.
    try:
        first_row = table.rows[0]
        trPr = first_row._tr.find(qn('w:trPr'))
        # Check if tblHeader element exists and is not explicitly disabled
        tblHeader_elem = trPr.find(qn('w:tblHeader')) if trPr is not None else None
        if tblHeader_elem is not None:
            val = tblHeader_elem.get(qn('w:val'))
            # tblHeader present with val=None, "true", "1", or "" all mean enabled
            # Only val="false" or val="0" would mean disabled
            if val is None or val.lower() not in ('false', '0'):
                print(f"PASS: Component 3 — Header row repeat is enabled (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — tblHeader element present but disabled (val={val})")
        else:
            print(f"FAIL: Component 3 — Header row repeat not enabled on first row")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Data rows populated (0.15 points)
    # Check that at least 25 of the 30 data rows have non-empty content in multiple columns
    try:
        populated_rows = 0
        for ri in range(1, len(table.rows)):
            cells = [cell.text.strip() for cell in table.rows[ri].cells]
            # A data row is populated if at least 3 of 5 cells have content
            non_empty = sum(1 for c in cells if c)
            if non_empty >= 3:
                populated_rows += 1

        if populated_rows >= 25:
            print(f"PASS: Component 4 — {populated_rows} data rows populated (0.15 pts)")
            total_score += 0.15
        elif populated_rows >= 10:
            partial = 0.08
            print(f"PARTIAL: Component 4 — Only {populated_rows} data rows populated, expected >= 25 ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {populated_rows} data rows populated, expected >= 25")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
