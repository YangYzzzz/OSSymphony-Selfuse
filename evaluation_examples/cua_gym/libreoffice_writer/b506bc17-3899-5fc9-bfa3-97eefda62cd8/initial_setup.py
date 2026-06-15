"""
Initial Setup: Contract document with 7 paragraphs, last paragraph at 12pt normal formatting.
Task ID: osworld_writer_strikethrough_last_para_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title / Contract heading
    heading = doc.add_paragraph()
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run("SERVICE AGREEMENT AND CONTRACT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"

    # Paragraph 2: Parties clause
    para2 = doc.add_paragraph()
    run2 = para2.add_run(
        "This Service Agreement (\"Agreement\") is entered into as of March 1, 2025, "
        "by and between Northgate Consulting LLC, a limited liability company organized "
        "under the laws of the State of Delaware (\"Service Provider\"), and Meridian "
        "Technologies Inc., a corporation incorporated in California (\"Client\")."
    )
    run2.font.size = Pt(12)
    run2.font.name = "Calibri"

    # Paragraph 3: Scope of Services
    para3 = doc.add_paragraph()
    run3a = para3.add_run("1. Scope of Services. ")
    run3a.bold = True
    run3a.font.size = Pt(12)
    run3a.font.name = "Calibri"
    run3b = para3.add_run(
        "Service Provider agrees to deliver software development, technical consulting, "
        "and project management services as outlined in Exhibit A attached hereto. "
        "All deliverables shall conform to the specifications mutually agreed upon "
        "in writing by both parties prior to commencement of each milestone."
    )
    run3b.font.size = Pt(12)
    run3b.font.name = "Calibri"

    # Paragraph 4: Term and Termination
    para4 = doc.add_paragraph()
    run4a = para4.add_run("2. Term and Termination. ")
    run4a.bold = True
    run4a.font.size = Pt(12)
    run4a.font.name = "Calibri"
    run4b = para4.add_run(
        "This Agreement shall commence on March 1, 2025 and continue through "
        "February 28, 2026, unless earlier terminated. Either party may terminate "
        "this Agreement with thirty (30) days' written notice. In the event of material "
        "breach, the non-breaching party may terminate immediately upon written notice."
    )
    run4b.font.size = Pt(12)
    run4b.font.name = "Calibri"

    # Paragraph 5: Compensation
    para5 = doc.add_paragraph()
    run5a = para5.add_run("3. Compensation. ")
    run5a.bold = True
    run5a.font.size = Pt(12)
    run5a.font.name = "Calibri"
    run5b = para5.add_run(
        "Client shall pay Service Provider a monthly retainer of $12,500 USD, "
        "invoiced on the first business day of each calendar month and due within "
        "fifteen (15) days of invoice receipt. Late payments shall accrue interest "
        "at a rate of 1.5% per month on the outstanding balance."
    )
    run5b.font.size = Pt(12)
    run5b.font.name = "Calibri"

    # Paragraph 6: Confidentiality
    para6 = doc.add_paragraph()
    run6a = para6.add_run("4. Confidentiality. ")
    run6a.bold = True
    run6a.font.size = Pt(12)
    run6a.font.name = "Calibri"
    run6b = para6.add_run(
        "Each party agrees to keep confidential all proprietary information, trade secrets, "
        "and non-public business information disclosed by the other party in connection with "
        "this Agreement. This obligation shall survive termination of the Agreement for a "
        "period of three (3) years. Disclosure is permitted only to employees with a need "
        "to know, subject to equivalent confidentiality obligations."
    )
    run6b.font.size = Pt(12)
    run6b.font.name = "Calibri"

    # Paragraph 7: Governing Law (last paragraph - clause being removed, normal formatting at 12pt)
    para7 = doc.add_paragraph()
    run7a = para7.add_run("5. Governing Law and Dispute Resolution. ")
    run7a.bold = True
    run7a.font.size = Pt(12)
    run7a.font.name = "Calibri"
    run7b = para7.add_run(
        "This Agreement shall be governed by and construed in accordance with the laws "
        "of the State of Delaware, without regard to its conflict of laws provisions. "
        "Any dispute arising under this Agreement shall be resolved through binding "
        "arbitration administered by the American Arbitration Association under its "
        "Commercial Arbitration Rules. The arbitration shall take place in Wilmington, "
        "Delaware, and the decision of the arbitrator shall be final and binding."
    )
    run7b.font.size = Pt(12)
    run7b.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
