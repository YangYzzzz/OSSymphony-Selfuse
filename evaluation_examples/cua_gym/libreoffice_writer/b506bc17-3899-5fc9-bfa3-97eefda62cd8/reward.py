"""
Reward Script: Apply strikethrough, gray color, and reduced font size to last paragraph
Task ID: osworld_writer_strikethrough_last_para_009
Domain: libreoffice_writer
Scoring:
  Component 1: Strikethrough applied to all runs in last paragraph (0.4 pts)
  Component 2: Font color changed to gray on all runs in last paragraph (0.3 pts)
  Component 3: Font size reduced by 2pt (12pt -> 10pt) on all runs in last paragraph (0.3 pts)
  Total: 1.0
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_009'


def color_distance(c1, c2):
    """Compute Euclidean RGB color distance."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Apply strikethrough, gray font color, and reduce font size by 2pts
    to all text in the last paragraph of the contract document.
    Initial state: last paragraph has 12pt font, no strikethrough, no special color.
    Golden state: last paragraph has strikethrough=True, color=gray(0x808080), size=10pt.
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document should have exactly 7 paragraphs
    if len(doc.paragraphs) == 0:
        print("CRITICAL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    last_para = doc.paragraphs[-1]
    text_runs = [run for run in last_para.runs if run.text.strip()]

    if not text_runs:
        print("FAIL: Last paragraph has no text runs")
        print(f"Score: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Component 1: Strikethrough applied to ALL text runs in last paragraph (0.4 points)
    # This verifies the primary formatting change — marking text as deleted
    try:
        all_strike = all(run.font.strike is True for run in text_runs)
        strike_count = sum(1 for run in text_runs if run.font.strike is True)
        if all_strike:
            print(f"PASS: Component 1 — Strikethrough applied to all {len(text_runs)} text run(s) in last paragraph (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Strikethrough not applied to all runs. {strike_count}/{len(text_runs)} runs have strikethrough.")
    except Exception as e:
        print(f"ERROR: Component 1 (strikethrough check) — {e}")

    # Component 2: Font color changed to gray on ALL text runs in last paragraph (0.3 points)
    # Gray is defined as RGBColor(0x80, 0x80, 0x80) = (128, 128, 128)
    # Allow some tolerance for near-gray colors (distance < 50 from target gray)
    GRAY_TARGET = (128, 128, 128)
    COLOR_TOLERANCE = 50
    try:
        gray_count = 0
        non_gray_details = []
        for run in text_runs:
            try:
                rgb = run.font.color.rgb
                dist = color_distance(tuple(rgb), GRAY_TARGET)
                if dist < COLOR_TOLERANCE:
                    gray_count += 1
                else:
                    non_gray_details.append(f"run text={repr(run.text[:30])}, color={rgb}, dist={dist:.1f}")
            except Exception:
                # color.rgb raises if no explicit color set (inherited)
                non_gray_details.append(f"run text={repr(run.text[:30])}, no explicit color set")

        if gray_count == len(text_runs):
            print(f"PASS: Component 2 — Gray font color applied to all {len(text_runs)} text run(s) in last paragraph (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Gray color not applied to all runs. {gray_count}/{len(text_runs)} runs have gray color.")
            for detail in non_gray_details:
                print(f"  Non-gray run: {detail}")
    except Exception as e:
        print(f"ERROR: Component 2 (color check) — {e}")

    # Component 3: Font size reduced to 10pt (from original 12pt) in ALL text runs (0.3 points)
    # Task says: reduce font size by 2 points. Initial size is 12pt, so target is 10pt.
    TARGET_SIZE_PT = 10.0
    try:
        correct_size_count = 0
        wrong_size_details = []
        for run in text_runs:
            if run.font.size is not None:
                actual_pt = run.font.size.pt
                if abs(actual_pt - TARGET_SIZE_PT) < 0.5:
                    correct_size_count += 1
                else:
                    wrong_size_details.append(f"run text={repr(run.text[:30])}, size={actual_pt}pt (expected {TARGET_SIZE_PT}pt)")
            else:
                wrong_size_details.append(f"run text={repr(run.text[:30])}, size=None (inherited, not explicitly set)")

        if correct_size_count == len(text_runs):
            print(f"PASS: Component 3 — Font size reduced to {TARGET_SIZE_PT}pt on all {len(text_runs)} text run(s) in last paragraph (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Font size not {TARGET_SIZE_PT}pt on all runs. {correct_size_count}/{len(text_runs)} runs have correct size.")
            for detail in wrong_size_details:
                print(f"  Wrong size run: {detail}")
    except Exception as e:
        print(f"ERROR: Component 3 (font size check) — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
