"""
Initial Setup: Create a Writer document with a Results section (no table yet)
Task ID: writer_acad_027
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('Effects of Cognitive Load on Working Memory Performance', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract-like intro paragraph
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing = 2.0
    run = para.add_run(
        'This study examined the effects of cognitive load on working memory performance '
        'across different age groups. Participants (N = 120) were randomly assigned to one '
        'of three experimental conditions: low cognitive load, moderate cognitive load, and '
        'high cognitive load. Each participant completed a series of memory tasks while '
        'simultaneously performing a secondary task calibrated to the assigned load level.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Method section heading
    method_heading = doc.add_heading('Method', level=2)

    para2 = doc.add_paragraph()
    para2.paragraph_format.line_spacing = 2.0
    run2 = para2.add_run(
        'Participants were recruited from an undergraduate psychology pool at a large '
        'Midwestern university. Inclusion criteria required normal or corrected-to-normal '
        'vision and no history of neurological disorders. The sample consisted of 68 females '
        'and 52 males, with a mean age of 21.4 years (SD = 3.2). All participants provided '
        'informed consent prior to testing.'
    )
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # Results section heading
    results_heading = doc.add_heading('Results', level=2)

    para3 = doc.add_paragraph()
    para3.paragraph_format.line_spacing = 2.0
    run3 = para3.add_run(
        'Descriptive statistics for all key variables are presented in the table below. '
        'A one-way ANOVA revealed significant differences among groups on the primary '
        'outcome measure of working memory accuracy, F(2, 117) = 14.83, p < .001, '
        '\u03b7\u00b2 = .20. Post hoc comparisons using Tukey\'s HSD indicated that the '
        'high cognitive load group performed significantly worse than both the low and '
        'moderate load groups.'
    )
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    # Discussion section heading (to show document continues)
    discussion_heading = doc.add_heading('Discussion', level=2)

    para4 = doc.add_paragraph()
    para4.paragraph_format.line_spacing = 2.0
    run4 = para4.add_run(
        'The results of this study support the hypothesis that increased cognitive load '
        'negatively impacts working memory performance. These findings are consistent with '
        'Baddeley\'s model of working memory, which posits that the central executive has '
        'limited capacity for processing concurrent tasks. The significant effect of load '
        'condition on accuracy scores suggests that the secondary task effectively consumed '
        'cognitive resources that would otherwise be available for memory encoding.'
    )
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
