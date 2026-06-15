"""
Reward Script: Create APA-style academic results table
Task ID: writer_acad_027
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with correct dimensions (6 rows x 4 cols)
  Component 2 (0.25): Header row contains correct column names
  Component 3 (0.25): 5 data rows with non-empty content
  Component 4 (0.25): APA-style borders (horizontal only on header top/bottom and last row bottom)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_027'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_cell_border(cell, border_name):
    """Get border value ('single', 'nil', 'none', etc.) for a cell's specified border."""
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return None
    border_el = tcBorders.find(qn(f'w:{border_name}'))
    if border_el is None:
        return None
    return border_el.get(qn('w:val'))


def is_visible_border(val):
    """Check if a border value represents a visible border."""
    if val is None:
        return False
    return val not in ('nil', 'none')


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions — 6 rows, 4 columns (0.25 points)
    try:
        if len(doc.tables) < 1:
            print(f"FAIL: Component 1 — No tables found in document (found {len(doc.tables)})")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 6 and num_cols == 4:
                print(f"PASS: Component 1 — Table has {num_rows} rows x {num_cols} cols (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — Expected 6x4 table, found {num_rows}x{num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Gate: must have at least one table for remaining checks
    if len(doc.tables) < 1:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 2: Header row contains correct column names (0.25 points)
    try:
        expected_headers = ['Variable', 'Mean', 'SD', 'p-value']
        if len(table.rows) > 0 and len(table.columns) >= 4:
            actual_headers = [table.cell(0, c).text.strip() for c in range(4)]
            matches = sum(1 for exp, act in zip(expected_headers, actual_headers)
                         if exp.lower() == act.lower())
            if matches == 4:
                print(f"PASS: Component 2 — Headers match: {actual_headers} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — Expected {expected_headers}, found {actual_headers} ({matches}/4 match)")
        else:
            print(f"FAIL: Component 2 — Table too small for header check")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 5 data rows with non-empty content (0.25 points)
    try:
        if len(table.rows) >= 6:
            filled_rows = 0
            for ri in range(1, 6):
                row_texts = [table.cell(ri, c).text.strip() for c in range(min(4, len(table.columns)))]
                if any(t for t in row_texts):
                    filled_rows += 1
            if filled_rows == 5:
                print(f"PASS: Component 3 — All 5 data rows have content (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — Only {filled_rows}/5 data rows have content")
        else:
            print(f"FAIL: Component 3 — Not enough rows for 5 data rows (has {len(table.rows)})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: APA-style borders (0.25 points)
    # Header row (row 0): top border visible, bottom border visible
    # Last data row (row 5): bottom border visible
    # No vertical borders anywhere (left/right on all cells should be nil/none)
    try:
        if len(table.rows) >= 6 and len(table.columns) >= 4:
            issues = []

            # Check header row top and bottom borders
            for ci in range(4):
                cell = table.cell(0, ci)
                top_val = get_cell_border(cell, 'top')
                bottom_val = get_cell_border(cell, 'bottom')
                if not is_visible_border(top_val):
                    issues.append(f"Row 0, Col {ci}: top border missing (val={top_val})")
                if not is_visible_border(bottom_val):
                    issues.append(f"Row 0, Col {ci}: bottom border missing (val={bottom_val})")

            # Check last row bottom border
            last_row_idx = len(table.rows) - 1
            for ci in range(4):
                cell = table.cell(last_row_idx, ci)
                bottom_val = get_cell_border(cell, 'bottom')
                if not is_visible_border(bottom_val):
                    issues.append(f"Row {last_row_idx}, Col {ci}: bottom border missing (val={bottom_val})")

            # Check that no vertical borders exist (left/right on all cells should be nil/none)
            vertical_visible = 0
            for ri in range(len(table.rows)):
                for ci in range(min(4, len(table.columns))):
                    cell = table.cell(ri, ci)
                    left_val = get_cell_border(cell, 'left')
                    right_val = get_cell_border(cell, 'right')
                    if is_visible_border(left_val):
                        vertical_visible += 1
                    if is_visible_border(right_val):
                        vertical_visible += 1

            if vertical_visible > 0:
                issues.append(f"{vertical_visible} vertical borders found (should be 0 for APA style)")

            if len(issues) == 0:
                print(f"PASS: Component 4 — APA borders correct: header top/bottom, last row bottom, no vertical (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 4 — Border issues: {'; '.join(issues[:5])}")
        else:
            print(f"FAIL: Component 4 — Table too small for border verification")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
