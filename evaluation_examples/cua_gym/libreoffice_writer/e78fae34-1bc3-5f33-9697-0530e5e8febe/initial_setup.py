"""
Initial Setup: Coding standards document with BEST PRACTICE tip boxes (no background color)
Task ID: writer_para_056
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_para_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1 - 'Company Coding Standards Guide'
    doc.add_heading('Company Coding Standards Guide', level=1)

    # Paragraph 2: Heading 2 - 'Naming Conventions'
    doc.add_heading('Naming Conventions', level=2)

    # Paragraph 3: Normal paragraph about camelCase, PascalCase, UPPER_SNAKE_CASE
    doc.add_paragraph(
        'All variable names should use camelCase notation. '
        'Class names must use PascalCase. '
        'Constants should be UPPER_SNAKE_CASE.'
    )

    # Paragraph 4: BEST PRACTICE tip box — NO background color in initial state
    doc.add_paragraph(
        'BEST PRACTICE: Use descriptive variable names that clearly indicate their purpose. '
        'Avoid single-letter variables except in loop counters and mathematical expressions.'
    )

    # Paragraph 5: Heading 2 - 'Error Handling'
    doc.add_heading('Error Handling', level=2)

    # Paragraph 6: Normal paragraph about error handling
    doc.add_paragraph(
        'All public methods must include proper error handling. '
        'Exceptions should be caught at the appropriate level and logged with sufficient context for debugging.'
    )

    # Paragraph 7: BEST PRACTICE tip box — NO background color in initial state
    doc.add_paragraph(
        'BEST PRACTICE: Create custom exception classes for domain-specific errors. '
        'Always include the original exception as a cause when re-throwing to preserve the stack trace.'
    )

    # Paragraph 8: Heading 2 - 'Code Documentation'
    doc.add_heading('Code Documentation', level=2)

    # Paragraph 9: Normal paragraph about code documentation
    doc.add_paragraph(
        'All public APIs must have comprehensive documentation including parameter descriptions, '
        'return values, and examples.'
    )

    # Paragraph 10: BEST PRACTICE tip box — NO background color in initial state
    doc.add_paragraph(
        'BEST PRACTICE: Write documentation as if the reader has no context about the code. '
        'Include code examples for complex methods and document any side effects.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
