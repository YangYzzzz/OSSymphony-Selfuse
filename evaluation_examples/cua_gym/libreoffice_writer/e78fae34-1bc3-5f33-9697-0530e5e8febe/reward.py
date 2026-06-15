"""
Reward Script: Set paragraph background color to light green (#E8F5E9) for BEST PRACTICE tip boxes
Task ID: writer_para_056
Domain: libreoffice_writer
Scoring:
  - Component 1: Para 4 (first BEST PRACTICE) has background color #E8F5E9 (0.34 pts)
  - Component 2: Para 7 (second BEST PRACTICE) has background color #E8F5E9 (0.33 pts)
  - Component 3: Para 10 (third BEST PRACTICE) has background color #E8F5E9 (0.33 pts)
  (No other paragraphs should have the target color — verified as precondition gate)
Total: 1.0
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_para_056'
TARGET_COLOR = 'E8F5E9'  # Light green as stored in XML (no '#' prefix)

# Indices (0-based) of BEST PRACTICE paragraphs that should have the background
BEST_PRACTICE_INDICES = [3, 6, 9]  # Para 4, 7, 10 (1-based) = indices 3, 6, 9
# Indices of all OTHER paragraphs that should NOT have any background
OTHER_INDICES = [0, 1, 2, 4, 5, 7, 8]  # Para 1, 2, 3, 5, 6, 8, 9 (1-based)


def get_paragraph_bg_color(para):
    """
    Extract the paragraph background fill color from the shading XML element.
    Returns the hex color string (e.g. 'E8F5E9') or None if no shading is applied.
    """
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    return fill


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — precondition gate
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: Check document has expected 10 paragraphs
    if len(doc.paragraphs) != 10:
        print(f"CRITICAL: Expected 10 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs

    # Component 1: Para 4 (index 3) — first BEST PRACTICE paragraph has bg=E8F5E9 (0.34 pts)
    try:
        para4 = paragraphs[3]
        bg4 = get_paragraph_bg_color(para4)
        if bg4 is not None and bg4.upper() == TARGET_COLOR:
            print(f"PASS: Component 1 — Para 4 has background color #{TARGET_COLOR} (0.34 pts)")
            print(f"      Text: {repr(para4.text[:60])}")
            total_score += 0.34
        else:
            print(f"FAIL: Component 1 — Para 4 expected background #{TARGET_COLOR}, found: {bg4}")
            print(f"      Text: {repr(para4.text[:60])}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Para 7 (index 6) — second BEST PRACTICE paragraph has bg=E8F5E9 (0.33 pts)
    try:
        para7 = paragraphs[6]
        bg7 = get_paragraph_bg_color(para7)
        if bg7 is not None and bg7.upper() == TARGET_COLOR:
            print(f"PASS: Component 2 — Para 7 has background color #{TARGET_COLOR} (0.33 pts)")
            print(f"      Text: {repr(para7.text[:60])}")
            total_score += 0.33
        else:
            print(f"FAIL: Component 2 — Para 7 expected background #{TARGET_COLOR}, found: {bg7}")
            print(f"      Text: {repr(para7.text[:60])}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Para 10 (index 9) — third BEST PRACTICE paragraph has bg=E8F5E9 (0.33 pts)
    try:
        para10 = paragraphs[9]
        bg10 = get_paragraph_bg_color(para10)
        if bg10 is not None and bg10.upper() == TARGET_COLOR:
            print(f"PASS: Component 3 — Para 10 has background color #{TARGET_COLOR} (0.33 pts)")
            print(f"      Text: {repr(para10.text[:60])}")
            total_score += 0.33
        else:
            print(f"FAIL: Component 3 — Para 10 expected background #{TARGET_COLOR}, found: {bg10}")
            print(f"      Text: {repr(para10.text[:60])}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
