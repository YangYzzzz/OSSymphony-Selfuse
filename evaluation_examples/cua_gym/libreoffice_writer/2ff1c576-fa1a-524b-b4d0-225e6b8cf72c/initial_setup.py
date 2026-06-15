"""
Initial Setup: Court opinion document with default styles only
Task ID: writer_legal_086
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_086'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Use only default Normal style throughout — no custom styles
    # This is a draft court opinion that needs formatting

    # Title area (just plain text, no special style)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("SUPREME COURT OF THE STATE OF COLUMBIA")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("No. 2025-CV-04817")
    run.font.size = Pt(11)

    doc.add_paragraph()  # blank line

    p = doc.add_paragraph()
    run = p.add_run("WESTFIELD NATIONAL BANK, Plaintiff-Appellant,")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("v.")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("THORNBERRY DEVELOPMENT GROUP, LLC, Defendant-Appellee.")
    run.font.size = Pt(11)

    doc.add_paragraph()  # blank line

    # Opinion heading — currently just plain text
    p = doc.add_paragraph()
    run = p.add_run("OPINION OF THE COURT")
    run.font.size = Pt(11)

    doc.add_paragraph()  # blank line

    # Subheading
    p = doc.add_paragraph()
    run = p.add_run("Factual and Procedural Background")
    run.font.size = Pt(11)

    # Body paragraphs — all in default style
    body_texts = [
        "This appeal arises from the Circuit Court of Meridian County's entry of summary judgment in favor of Thornberry Development Group, LLC on claims of breach of a commercial loan agreement and fraudulent misrepresentation. The underlying dispute concerns a $4.2 million construction loan originated in March 2023 for the development of a mixed-use commercial property located at 1847 Crescent Avenue, Millbrook.",
        "Westfield National Bank extended the loan pursuant to a Master Loan Agreement dated March 15, 2023, which incorporated standard draw-down provisions contingent upon verified completion of specified construction milestones. Between April 2023 and November 2023, Thornberry submitted seven draw requests totaling $3.1 million, each accompanied by certifications from its project manager, Elena Vasquez, attesting to the completion of designated phases.",
        "In December 2023, an independent inspection commissioned by Westfield revealed that construction progress had been materially overstated in at least four of the seven draw certifications. The inspection report, prepared by Hargrove Engineering Associates, estimated that actual completed work corresponded to approximately $1.8 million in value, creating an apparent discrepancy of $1.3 million between disbursed funds and verified progress.",
        "Westfield commenced this action on February 12, 2024, asserting claims for breach of contract, fraudulent misrepresentation, and unjust enrichment. Following discovery, both parties moved for summary judgment. The Circuit Court granted Thornberry's motion on all counts, finding that the draw certifications were opinions rather than statements of fact and that Westfield had failed to exercise its contractual right of independent verification prior to disbursement.",
    ]

    for text in body_texts:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)

    # Another subheading
    p = doc.add_paragraph()
    run = p.add_run("Standard of Review")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("We review a grant of summary judgment de novo, applying the same standard as the trial court. Henderson v. Pacific Mutual Insurance Co., 298 Col. 415, 420 (2019). Summary judgment is appropriate when the pleadings, depositions, answers to interrogatories, and admissions on file, together with the affidavits, if any, show that there is no genuine issue as to any material fact and that the moving party is entitled to judgment as a matter of law. Col. R. Civ. P. 56(c).")
    run.font.size = Pt(11)

    # Another subheading
    p = doc.add_paragraph()
    run = p.add_run("Analysis")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("The central question before this Court is whether the draw certifications submitted by Thornberry constitute actionable representations of fact or mere expressions of opinion. This distinction is critical because, under Columbia law, a claim of fraudulent misrepresentation requires proof of a false statement of material fact. Brennan v. Consolidated Financial Services, 312 Col. 89, 97 (2021).")
    run.font.size = Pt(11)

    # Block quotation — currently just normal paragraph
    p = doc.add_paragraph()
    run = p.add_run('As this Court observed in Brennan: "A statement of opinion, even if erroneous, does not give rise to liability for fraud unless the speaker purports to have special knowledge of facts that are not available to the other party and that contradict the opinion expressed." 312 Col. at 98.')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("We find this principle distinguishable from the present case. Unlike a general opinion about market conditions or future prospects, a draw certification is a specific factual representation that designated construction milestones have been completed. The Master Loan Agreement expressly characterized these certifications as statements of fact upon which the lender was entitled to rely.")
    run.font.size = Pt(11)

    # Another subheading
    p = doc.add_paragraph()
    run = p.add_run("Conclusion")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("For the foregoing reasons, the judgment of the Circuit Court is reversed and the case is remanded for further proceedings consistent with this opinion.")
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("DATED this 15th day of January, 2025.")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Chief Justice Margaret A. Thornton")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Associate Justice David R. Keane, concurring")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Associate Justice Priya S. Chakraborty, concurring")
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
