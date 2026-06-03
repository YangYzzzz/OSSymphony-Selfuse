"""
Initial Setup: Payroll Procedures Manual - raw procedural text without tables
Task ID: writer_hr_090
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_090'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Payroll_Procedures_Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    subtitle.add_run('\n')
    run2 = subtitle.add_run('Human Resources Department - Payroll Division')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()  # spacing

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This Payroll Procedures Manual serves as the comprehensive reference guide for all '
        'payroll processing activities within Meridian Technologies Inc. The procedures outlined '
        'herein are designed to ensure accurate, timely, and compliant payroll processing for all '
        'employees across our domestic operations.'
    )
    doc.add_paragraph(
        'All payroll staff members are expected to follow these procedures strictly. Deviations '
        'from established protocols must be approved by the Payroll Manager or the Director of '
        'Human Resources. This manual is reviewed and updated annually to reflect changes in '
        'federal and state tax regulations, company policies, and best practices.'
    )
    doc.add_paragraph(
        'Meridian Technologies operates on a bi-weekly payroll cycle with 26 pay periods per '
        'calendar year. The payroll calendar for 2026 details each pay period, including the '
        'period start date, period end date, timesheet deadline, and pay date. Refer to the '
        'payroll calendar section for complete scheduling information.'
    )

    # --- Section 2: Payroll Calendar ---
    doc.add_heading('2. Payroll Calendar for 2026', level=1)
    doc.add_paragraph(
        'The 2026 payroll calendar follows a bi-weekly schedule with 26 pay periods. Each pay '
        'period spans 14 calendar days. Timesheets must be submitted by 5:00 PM on the Monday '
        'following the end of each pay period. Paychecks are issued on the Friday following '
        'timesheet submission. The payroll calendar should be distributed to all department '
        'managers at the beginning of each fiscal year.'
    )
    doc.add_paragraph(
        'Note: When a pay date falls on a federal holiday, paychecks will be issued on the '
        'preceding business day. Employees enrolled in direct deposit will see funds available '
        'by 6:00 AM on the pay date.'
    )

    # --- Section 3: Tax Withholding ---
    doc.add_heading('3. Tax Withholding Reference', level=1)
    doc.add_paragraph(
        'Federal income tax withholding is calculated based on the employee\'s W-4 filing status '
        'and the IRS Publication 15-T tax tables. Meridian Technologies withholds federal income '
        'tax, Social Security tax (6.2%), and Medicare tax (1.45%) from each paycheck. Employees '
        'earning above the Social Security wage base ($168,600 for 2026) are exempt from '
        'additional Social Security withholding for the remainder of the year.'
    )
    doc.add_paragraph(
        'State income tax withholding varies by the employee\'s state of residence. Meridian '
        'Technologies currently has employees in California, New York, Texas, Florida, Illinois, '
        'and Washington. Texas, Florida, and Washington have no state income tax. Refer to the '
        'tax withholding reference tables for applicable rates by state and filing status.'
    )

    # --- Section 4: Deduction Codes ---
    doc.add_heading('4. Deduction Codes and Descriptions', level=1)
    doc.add_paragraph(
        'The payroll system uses standardized deduction codes to track all voluntary and '
        'involuntary deductions from employee paychecks. Each code is associated with a specific '
        'deduction type, description, and processing priority. Payroll analysts must use the '
        'correct deduction code when setting up new deductions or modifying existing ones. '
        'A complete listing of all active deduction codes should be referenced when processing '
        'payroll adjustments.'
    )
    doc.add_paragraph(
        'Pre-tax deductions (codes beginning with "PT") reduce taxable income and are processed '
        'before tax calculations. Post-tax deductions (codes beginning with "AT") are taken after '
        'all tax withholdings. Garnishment codes (beginning with "GR") follow federal and state '
        'priority rules as outlined in the garnishment processing section.'
    )

    # --- Section 5: Overtime Calculations ---
    doc.add_heading('5. Overtime Calculation Procedures', level=1)
    doc.add_paragraph(
        'Overtime compensation is governed by the Fair Labor Standards Act (FLSA) and applicable '
        'state labor laws. Non-exempt employees are entitled to overtime pay at 1.5 times their '
        'regular hourly rate for hours worked in excess of 40 hours per workweek. California-based '
        'employees are additionally entitled to daily overtime for hours exceeding 8 in a single '
        'workday and double-time for hours exceeding 12.'
    )
    doc.add_paragraph(
        'The payroll system automatically calculates overtime based on approved timesheets. '
        'However, payroll analysts should verify calculations for employees with multiple pay '
        'rates, shift differentials, or bonus payments that affect the regular rate of pay. '
        'Example calculations for common overtime scenarios should be reviewed to ensure '
        'compliance with FLSA weighted average overtime rules.'
    )

    # --- Section 6: Garnishment Processing ---
    doc.add_heading('6. Garnishment Processing Procedures', level=1)
    doc.add_paragraph(
        'Garnishment orders must be processed within 7 business days of receipt by the Payroll '
        'Department. All garnishment orders are logged in the Garnishment Tracking System and '
        'assigned to a payroll analyst for processing. Federal law establishes a priority order '
        'for garnishments when multiple orders exist for a single employee.'
    )
    doc.add_paragraph(
        'The garnishment priority order determines which obligations are satisfied first when '
        'disposable earnings are insufficient to cover all garnishments. Child support orders '
        'receive the highest priority under federal law, followed by federal tax levies, federal '
        'student loan garnishments, creditor garnishments, and state tax levies. The maximum '
        'withholding limits vary by garnishment type and the employee\'s dependency status.'
    )

    # --- Section 7: Year-End Processing ---
    doc.add_heading('7. Year-End Processing', level=1)
    doc.add_paragraph(
        'Year-end processing is a critical phase of the payroll cycle that requires careful '
        'coordination across the Payroll, Human Resources, and Finance departments. The '
        'year-end processing checklist must be completed in sequence to ensure accurate W-2 '
        'generation and regulatory compliance.'
    )
    doc.add_paragraph(
        'Key year-end activities include: verifying employee demographic data, reconciling '
        'quarterly tax filings (Forms 941), reviewing benefit deduction totals, processing '
        'final adjustments, generating W-2 forms, filing federal and state tax returns, and '
        'updating the payroll system for the new tax year. All year-end tasks must be completed '
        'by January 31st of the following year.'
    )

    # --- Section 8: Error Correction ---
    doc.add_heading('8. Error Correction Procedures', level=1)
    doc.add_paragraph(
        'Payroll errors must be identified, documented, and corrected promptly. The error '
        'correction process follows a structured workflow to ensure accountability and prevent '
        'recurrence. Errors are classified into three categories based on severity and impact.'
    )
    doc.add_paragraph(
        'Category 1 (Critical): Errors affecting gross pay by more than $500 or impacting tax '
        'withholdings. These require immediate correction in the next available pay cycle and '
        'notification to the affected employee within 24 hours. Category 2 (Moderate): Errors '
        'affecting gross pay between $50 and $500. These are corrected in the next regular pay '
        'cycle. Category 3 (Minor): Errors affecting gross pay by less than $50 or cosmetic '
        'errors. These are batched and corrected during the next scheduled payroll run.'
    )
    doc.add_paragraph(
        'The error correction flowchart outlines the decision-making process from error '
        'identification through resolution. Each step requires documentation in the Payroll '
        'Error Log, including the error type, root cause, corrective action taken, and '
        'preventive measures implemented.'
    )

    # --- Section 9: Document Control ---
    doc.add_heading('9. Document Control', level=1)
    doc.add_paragraph(
        'This manual is maintained by the Payroll Manager and reviewed annually. Version history '
        'and revision dates are tracked in the document control log. Any questions regarding '
        'payroll procedures should be directed to the Payroll Department at ext. 4500 or '
        'payroll@meridiantech.com.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
