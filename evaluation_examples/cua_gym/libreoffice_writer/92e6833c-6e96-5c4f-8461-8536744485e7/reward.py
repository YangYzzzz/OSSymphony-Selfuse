"""
Reward Script: Payroll Procedures Manual with tables and cross-references
Task ID: writer_hr_090
Domain: libreoffice_writer
Scoring:
  C1 (0.20) - Payroll calendar table (26 bi-weekly pay periods)
  C2 (0.15) - Tax withholding reference tables (Federal + State)
  C3 (0.15) - Deduction codes table (15+ entries)
  C4 (0.15) - Overtime calculation tables
  C5 (0.10) - Garnishment priority table
  C6 (0.10) - Year-end checklist + error correction tables
  C7 (0.15) - Table numbering and cross-references in text
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_090'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    tables = doc.tables
    num_tables = len(tables)
    print(f"INFO: Document has {len(doc.paragraphs)} paragraphs and {num_tables} tables")

    # Precondition: must have tables at all (initial has 0)
    if num_tables == 0:
        print("FAIL: No tables found in document. Task requires multiple tables.")
        print("REWARD: 0.0")
        return 0.0

    # =========================================================================
    # Component 1: Payroll Calendar Table (0.20 points)
    # Must have a table with ~26 data rows for bi-weekly pay periods
    # =========================================================================
    try:
        payroll_cal_found = False
        for t in tables:
            rows = len(t.rows)
            cols = len(t.columns)
            if rows < 20 or cols < 3:
                continue
            # Check if header suggests payroll calendar
            header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
            if any(kw in header_text for kw in ['pay period', 'pay date', 'period start', 'period end']):
                # Count data rows (non-header)
                data_rows = rows - 1
                if data_rows >= 24:  # Allow some tolerance around 26
                    payroll_cal_found = True
                    print(f"PASS: Component 1 - Payroll calendar table found: {rows} rows x {cols} cols, {data_rows} pay periods (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"PARTIAL: Component 1 - Payroll calendar table found but only {data_rows} periods (need >= 24)")
                    total_score += 0.10
                    payroll_cal_found = True
                break
        if not payroll_cal_found:
            print("FAIL: Component 1 - No payroll calendar table with pay period rows found")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # =========================================================================
    # Component 2: Tax Withholding Reference Tables (0.15 points)
    # Need at least 2 tables: federal tax brackets + state tax rates
    # =========================================================================
    try:
        tax_tables_found = 0
        for t in tables:
            header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
            # Federal tax table: look for bracket/rate/filer columns
            if any(kw in header_text for kw in ['tax bracket', 'single filer', 'married', 'filing']):
                tax_tables_found += 1
            # State tax table: look for state/rate columns
            elif any(kw in header_text for kw in ['state', 'tax type', 'rate range']):
                if 'state' in header_text and ('rate' in header_text or 'tax' in header_text):
                    tax_tables_found += 1

        if tax_tables_found >= 2:
            print(f"PASS: Component 2 - {tax_tables_found} tax reference tables found (0.15 pts)")
            total_score += 0.15
        elif tax_tables_found == 1:
            print(f"PARTIAL: Component 2 - Only {tax_tables_found} tax table found (need 2)")
            total_score += 0.07
        else:
            print("FAIL: Component 2 - No tax withholding reference tables found")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # =========================================================================
    # Component 3: Deduction Codes Table with 15+ entries (0.15 points)
    # =========================================================================
    try:
        deduction_table_found = False
        for t in tables:
            header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
            if any(kw in header_text for kw in ['code', 'deduction']) and 'description' in header_text:
                data_rows = len(t.rows) - 1
                if data_rows >= 15:
                    deduction_table_found = True
                    print(f"PASS: Component 3 - Deduction codes table with {data_rows} entries (0.15 pts)")
                    total_score += 0.15
                elif data_rows >= 10:
                    deduction_table_found = True
                    print(f"PARTIAL: Component 3 - Deduction codes table found but only {data_rows} entries (need 15+)")
                    total_score += 0.07
                else:
                    deduction_table_found = True
                    print(f"FAIL: Component 3 - Deduction codes table found but only {data_rows} entries (need 15+)")
                break
        if not deduction_table_found:
            print("FAIL: Component 3 - No deduction codes table found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # =========================================================================
    # Component 4: Overtime Calculation Tables (0.15 points)
    # At least one table with overtime/OT calculation examples
    # =========================================================================
    try:
        ot_tables_found = 0
        for t in tables:
            header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
            all_text = ' '.join(cell.text.strip().lower() for row in t.rows for cell in row.cells)
            # Look for OT-related columns
            if any(kw in header_text for kw in ['ot hours', 'ot rate', 'overtime', 'hourly rate', 'regular hours']):
                ot_tables_found += 1
            elif 'component' in header_text and ('rate' in header_text or 'earnings' in header_text):
                # Weighted average OT table
                if 'ot' in all_text or 'overtime' in all_text:
                    ot_tables_found += 1

        if ot_tables_found >= 2:
            print(f"PASS: Component 4 - {ot_tables_found} overtime calculation tables found (0.15 pts)")
            total_score += 0.15
        elif ot_tables_found == 1:
            print(f"PARTIAL: Component 4 - {ot_tables_found} overtime table found (ideal: 2)")
            total_score += 0.10
        else:
            print("FAIL: Component 4 - No overtime calculation tables found")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # =========================================================================
    # Component 5: Garnishment Priority Table (0.10 points)
    # =========================================================================
    try:
        garnishment_found = False
        for t in tables:
            header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
            all_text = ' '.join(cell.text.strip().lower() for row in t.rows for cell in row.cells)
            if ('priority' in header_text or 'garnishment' in header_text) and \
               ('garnishment' in all_text or 'child support' in all_text or 'withholding' in header_text):
                garnishment_found = True
                print(f"PASS: Component 5 - Garnishment priority table found ({len(t.rows)} rows) (0.10 pts)")
                total_score += 0.10
                break
        if not garnishment_found:
            print("FAIL: Component 5 - No garnishment priority table found")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # =========================================================================
    # Component 6: Year-End Checklist + Error Correction Tables (0.10 points)
    # =========================================================================
    try:
        yearend_found = False
        error_found = False
        for t in tables:
            header_text = ' '.join(cell.text.strip().lower() for cell in t.rows[0].cells)
            all_text = ' '.join(cell.text.strip().lower() for row in t.rows[:3] for cell in row.cells)
            # Year-end checklist: look for step/task/deadline columns
            if ('task' in header_text or 'step' in header_text) and \
               ('deadline' in header_text or 'responsible' in header_text):
                if 'verify' in all_text or 'reconcil' in all_text or 'w-2' in all_text or 'year' in all_text:
                    yearend_found = True
            # Error correction flowchart: look for step/action/decision
            if ('step' in header_text) and ('action' in header_text or 'decision' in header_text):
                if 'error' in all_text or 'correction' in all_text:
                    error_found = True

        c6_score = 0.0
        if yearend_found:
            c6_score += 0.05
        if error_found:
            c6_score += 0.05
        if c6_score > 0:
            total_score += c6_score
            parts = []
            if yearend_found:
                parts.append("year-end checklist")
            if error_found:
                parts.append("error correction flowchart")
            print(f"PASS: Component 6 - Found: {', '.join(parts)} ({c6_score} pts)")
        else:
            print("FAIL: Component 6 - No year-end checklist or error correction tables found")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    # =========================================================================
    # Component 7: Table Numbering and Cross-References (0.15 points)
    # Tables should be labeled "Table N:" and referenced in body text
    # =========================================================================
    try:
        # Find "Table N" labels in paragraphs (caption-style: "Table 1: ...")
        table_labels = set()
        table_refs_in_text = set()
        for p in doc.paragraphs:
            text = p.text.strip()
            # Table labels: "Table N: description" at start of paragraph
            label_match = re.match(r'^Table\s+(\d+)\s*[:.]', text)
            if label_match:
                table_labels.add(int(label_match.group(1)))
            # Cross-references: "Table N" or "(see Table N)" in body text
            ref_matches = re.findall(r'Table\s+(\d+)', text)
            for m in ref_matches:
                # Only count as cross-ref if not a label line itself
                if not label_match:
                    table_refs_in_text.add(int(m))

        num_labels = len(table_labels)
        num_refs = len(table_refs_in_text)
        c7_score = 0.0

        # Need at least 5 labeled tables and cross-references to distinct tables
        if num_labels >= 5:
            c7_score += 0.08
        elif num_labels >= 3:
            c7_score += 0.04

        if num_refs >= 5:
            c7_score += 0.07
        elif num_refs >= 3:
            c7_score += 0.03

        if c7_score > 0:
            total_score += c7_score
            print(f"PASS: Component 7 - {num_labels} table labels, {num_refs} distinct tables cross-referenced ({c7_score} pts)")
        else:
            print(f"FAIL: Component 7 - Labels: {num_labels}, Cross-refs: {num_refs}")
    except Exception as e:
        print(f"ERROR: Component 7 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
