"""
Reward Script: PTO Tracking Table for 2026
Task ID: writer_hr_042
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25) - Table exists with correct dimensions (14 rows x 4 cols)
  Component 2 (0.20) - Header row has correct column names
  Component 3 (0.20) - Month names Jan-Dec in rows 1-12
  Component 4 (0.20) - Days Available = 1.25 for each month
  Component 5 (0.15) - Total row: label "Total" and Days Available = 15
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_042'

EXPECTED_MONTHS = [
    'January', 'February', 'March', 'April', 'May', 'June',
    'July', 'August', 'September', 'October', 'November', 'December'
]

EXPECTED_HEADERS = ['Month', 'Days Available', 'Days Used', 'Days Remaining']


def verify_task(file_path):
    """
    Verify PTO tracking table creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document. Task requires creating a PTO table.")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table exists with correct dimensions 14 rows x 4 cols (0.25 points)
    try:
        if num_rows == 14 and num_cols == 4:
            print(f"PASS: Component 1 - Table is 14x4 (header + 12 months + total) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - Expected 14 rows x 4 cols, found {num_rows} rows x {num_cols} cols")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Header row has correct column names (0.20 points)
    try:
        headers = [table.cell(0, c).text.strip() for c in range(min(num_cols, 4))]
        # Check headers case-insensitively for robustness
        headers_lower = [h.lower() for h in headers]
        expected_lower = [h.lower() for h in EXPECTED_HEADERS]
        if headers_lower == expected_lower:
            print(f"PASS: Component 2 - Headers match: {headers} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 - Expected headers {EXPECTED_HEADERS}, found {headers}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Month names Jan-Dec in rows 1-12 (0.20 points)
    try:
        if num_rows >= 13:
            actual_months = [table.cell(r, 0).text.strip() for r in range(1, 13)]
            actual_lower = [m.lower() for m in actual_months]
            expected_lower = [m.lower() for m in EXPECTED_MONTHS]
            matching = sum(1 for a, e in zip(actual_lower, expected_lower) if a == e)
            if matching == 12:
                print(f"PASS: Component 3 - All 12 month names correct (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 - {matching}/12 months match. Found: {actual_months}")
        else:
            print(f"FAIL: Component 3 - Not enough rows ({num_rows}) for 12 months")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Days Available = 1.25 for each month row (0.20 points)
    try:
        if num_rows >= 13 and num_cols >= 2:
            correct_count = 0
            for r in range(1, 13):
                val = table.cell(r, 1).text.strip()
                try:
                    if abs(float(val) - 1.25) < 0.001:
                        correct_count += 1
                except (ValueError, TypeError):
                    pass
            if correct_count == 12:
                print(f"PASS: Component 4 - All 12 months have Days Available = 1.25 (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 - {correct_count}/12 months have Days Available = 1.25")
        else:
            print(f"FAIL: Component 4 - Table too small for Days Available check")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Total row with label and sum (0.15 points)
    try:
        if num_rows >= 14:
            last_row = num_rows - 1
            total_label = table.cell(last_row, 0).text.strip().lower()
            total_val_text = table.cell(last_row, 1).text.strip()

            label_ok = 'total' in total_label
            try:
                total_val = float(total_val_text)
                val_ok = abs(total_val - 15.0) < 0.01
            except (ValueError, TypeError):
                val_ok = False

            if label_ok and val_ok:
                print(f"PASS: Component 5 - Total row has 'Total' label and sum = 15 (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 5 - Total row: label='{table.cell(last_row, 0).text.strip()}' (ok={label_ok}), value='{total_val_text}' (ok={val_ok})")
        else:
            print(f"FAIL: Component 5 - Not enough rows for total row")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
