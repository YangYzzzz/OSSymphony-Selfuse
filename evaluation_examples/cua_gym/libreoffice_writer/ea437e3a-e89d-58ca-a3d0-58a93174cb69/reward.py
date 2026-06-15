"""
Reward Script: Undo the last action performed on this document.
Task ID: writer_edit_023
Domain: libreoffice_writer
Scoring:
  Component 1: Marketing Q1 Budget reverts to '$50,000' (0.6 pts)
  Component 2: Marketing Annual Total reverts to '$105,000' (0.4 pts)
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_edit_023'
FILE_PATH = f'{WORKDIR}/budget_sheet.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Task: Undo the last action on budget_sheet.docx.
    The last action was changing Marketing Q1 Budget from '$50,000' to '$75,000'.
    After undo, Marketing Q1 Budget should be '$50,000' and Annual Total should be '$105,000'.

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — critical precondition
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition check: document has exactly 1 table with expected structure
    try:
        if len(doc.tables) < 1:
            print("CRITICAL: Document has no tables — unexpected structure")
            print("REWARD: 0.0")
            return 0.0
        table = doc.tables[0]
        if len(table.rows) < 3:
            print("CRITICAL: Table has fewer than 3 rows — unexpected structure")
            print("REWARD: 0.0")
            return 0.0
    except Exception as e:
        print(f"CRITICAL: Table access failed: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the Marketing row (it should be row index 2 based on initial setup,
    # but we search by label for robustness)
    marketing_row = None
    try:
        for row in table.rows:
            cells = row.cells
            if cells[0].text.strip().lower() == 'marketing':
                marketing_row = row
                break
    except Exception as e:
        print(f"ERROR: Cannot search for Marketing row: {e}")

    if marketing_row is None:
        print("FAIL: 'Marketing' row not found in table")
        print(f"Score: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Component 1: Marketing Q1 Budget reverts to '$50,000' (0.6 points)
    # Initial_env has '$75,000' (the bad edit), golden_env should have '$50,000' (after undo)
    try:
        q1_value = marketing_row.cells[1].text.strip()
        expected_q1 = '$50,000'
        if q1_value == expected_q1:
            print(f"PASS: Component 1 — Marketing Q1 Budget is '{q1_value}' (expected '{expected_q1}') (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 — Marketing Q1 Budget is '{q1_value}', expected '{expected_q1}'")
    except Exception as e:
        print(f"ERROR: Component 1 — Cannot read Marketing Q1 Budget: {e}")

    # Component 2: Marketing Annual Total reverts to '$105,000' (0.4 points)
    # Initial_env has '$130,000' (reflecting the bad edit), golden_env should have '$105,000' (after undo)
    try:
        annual_value = marketing_row.cells[3].text.strip()
        expected_annual = '$105,000'
        if annual_value == expected_annual:
            print(f"PASS: Component 2 — Marketing Annual Total is '{annual_value}' (expected '{expected_annual}') (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Marketing Annual Total is '{annual_value}', expected '{expected_annual}'")
    except Exception as e:
        print(f"ERROR: Component 2 — Cannot read Marketing Annual Total: {e}")

    # Supplementary check: other rows unchanged (informational only, not scored)
    # This verifies the undo did not corrupt unrelated rows
    try:
        expected_others = {
            'Engineering': ('$120,000', '$130,000', '$250,000'),
            'Human Resources': ('$40,000', '$42,000', '$82,000'),
            'Finance': ('$35,000', '$38,000', '$73,000'),
            'Operations': ('$95,000', '$100,000', '$195,000'),
            'Sales': ('$85,000', '$90,000', '$175,000'),
            'IT Support': ('$60,000', '$65,000', '$125,000'),
        }
        changed_rows = []
        for row in table.rows:
            dept = row.cells[0].text.strip()
            if dept in expected_others:
                exp_q1, exp_q2, exp_annual = expected_others[dept]
                actual_q1 = row.cells[1].text.strip()
                actual_q2 = row.cells[2].text.strip()
                actual_ann = row.cells[3].text.strip()
                if actual_q1 != exp_q1 or actual_q2 != exp_q2 or actual_ann != exp_annual:
                    changed_rows.append(dept)
                    print(f"INFO: Row '{dept}' values changed unexpectedly")
        if not changed_rows:
            print("INFO: All other department rows are unchanged (as expected)")
    except Exception as e:
        print(f"INFO: Could not verify other rows: {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path in VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
