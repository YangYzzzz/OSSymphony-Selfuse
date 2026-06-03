"""
Initial Setup: Undo the last action performed on the document.
Task ID: writer_edit_023
Domain: libreoffice_writer

The initial state has the document with Marketing row showing $75,000
(the result of the last action that changed it from $50,000 to $75,000).
The agent must undo (Ctrl+Z) to revert Marketing back to $50,000.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_023'
OUTPUT = f'{WORKDIR}/Desktop/budget_sheet.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Document title
    title = doc.add_heading('Annual Department Budget', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle paragraph
    subtitle = doc.add_paragraph('Fiscal Year 2024-2025')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacer

    # Budget table: headers + department rows
    # Columns: Department | Q1 Budget | Q2 Budget | Annual Total
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['Department', 'Q1 Budget', 'Q2 Budget', 'Annual Total']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    # Department data rows
    # NOTE: Marketing row shows $75,000 — this is the "last action" change
    # (original was $50,000, changed to $75,000; agent must undo to get $50,000 back)
    dept_data = [
        ['Engineering',   '$120,000', '$130,000', '$250,000'],
        ['Marketing',     '$75,000',  '$55,000',  '$130,000'],
        ['Human Resources','$40,000', '$42,000',  '$82,000' ],
        ['Finance',       '$35,000',  '$38,000',  '$73,000' ],
        ['Operations',    '$95,000',  '$100,000', '$195,000'],
        ['Sales',         '$85,000',  '$90,000',  '$175,000'],
        ['IT Support',    '$60,000',  '$65,000',  '$125,000'],
    ]

    for row_data in dept_data:
        row = table.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val

    doc.add_paragraph('')  # spacer

    # Footer note
    note = doc.add_paragraph(
        'Note: All figures are in USD. Budget subject to board approval.'
    )
    note.paragraph_format.space_before = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
