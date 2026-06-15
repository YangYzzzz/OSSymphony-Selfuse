"""
Initial Setup: Create a document with header containing 'Page ' followed by a page number field.
Task ID: writer_frd_060
Domain: libreoffice_writer

The document has a header with 'Page ' + PAGE field. The header text is selected
so the user can then save it as AutoText via Tools > AutoText.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_060'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Header with "Page " + page number field ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""  # Clear default

    # Add "Page " text
    run_text = hp.add_run("Page ")
    run_text.font.size = Pt(11)

    # Add PAGE field code (begin + instrText + end)
    run_begin = hp.add_run()
    fld_begin = run_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_begin._element.append(fld_begin)

    run_instr = hp.add_run()
    instr_text = run_instr._element.makeelement(qn('w:instrText'), {})
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    run_instr._element.append(instr_text)

    run_sep = hp.add_run()
    fld_sep = run_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run_sep._element.append(fld_sep)

    run_num = hp.add_run("1")  # Cached display value
    run_num.font.size = Pt(11)

    run_end = hp.add_run()
    fld_end = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_end._element.append(fld_end)

    # --- Body content: A realistic multi-page document ---
    doc.add_heading("Quarterly Financial Report - Q1 2025", level=1)

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "This report provides a comprehensive overview of the company's financial "
        "performance during the first quarter of 2025. Revenue grew by 12.3% compared "
        "to Q4 2024, driven primarily by strong demand in the enterprise segment and "
        "successful expansion into the Asia-Pacific market."
    )

    doc.add_heading("Revenue Breakdown", level=2)
    doc.add_paragraph(
        "Total revenue for Q1 2025 reached $48.7 million, representing a year-over-year "
        "increase of 18.5%. The enterprise division contributed $31.2 million, while the "
        "consumer segment generated $17.5 million. Key growth drivers included the launch "
        "of our new cloud-based analytics platform and a 25% increase in subscription renewals."
    )

    # Add a table for revenue data
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Division", "Q1 2025 ($M)", "Q4 2024 ($M)", "Change (%)"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ["Enterprise Cloud", "18.4", "16.1", "+14.3%"],
        ["Enterprise On-Prem", "12.8", "12.0", "+6.7%"],
        ["Consumer Premium", "10.2", "8.9", "+14.6%"],
        ["Consumer Free Tier", "7.3", "6.4", "+14.1%"],
        ["Total", "48.7", "43.4", "+12.2%"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacing

    doc.add_heading("Operating Expenses", level=2)
    doc.add_paragraph(
        "Operating expenses totaled $32.1 million in Q1 2025. Research and development "
        "spending increased to $14.8 million as we invested in next-generation AI capabilities. "
        "Sales and marketing expenses were $11.3 million, reflecting our continued investment "
        "in brand awareness and customer acquisition. General and administrative costs "
        "remained stable at $6.0 million."
    )

    # Page break to create a multi-page document
    doc.add_page_break()

    doc.add_heading("Regional Performance", level=2)
    doc.add_paragraph(
        "North America remained our largest market with $28.9 million in revenue, "
        "accounting for 59.3% of total revenue. Europe generated $12.4 million, "
        "showing steady growth of 8.2%. The Asia-Pacific region was the standout performer "
        "with $7.4 million in revenue, a 42.1% increase driven by our new Tokyo and "
        "Singapore offices."
    )

    doc.add_heading("Outlook for Q2 2025", level=2)
    doc.add_paragraph(
        "We project Q2 2025 revenue between $51 million and $54 million. This guidance "
        "reflects the anticipated launch of three new product features and the full-quarter "
        "impact of our recent partnership with Meridian Technologies. We expect operating "
        "margins to improve by 1-2 percentage points as efficiency programs take effect."
    )

    doc.add_heading("Key Risks and Considerations", level=2)
    doc.add_paragraph(
        "Management has identified several risk factors that could impact projections: "
        "ongoing supply chain disruptions, potential regulatory changes in the EU data "
        "privacy landscape, and increased competition from established players entering "
        "our market segment. Mitigation strategies are outlined in the appendix."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # --- GUI launch ---
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
