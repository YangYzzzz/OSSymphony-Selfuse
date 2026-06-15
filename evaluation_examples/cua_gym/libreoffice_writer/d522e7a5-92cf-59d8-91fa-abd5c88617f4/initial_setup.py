"""
Initial Setup: Create edit_specs.docx and product_photo.png on the Desktop
Task ID: osworld_multi_apps_writer_to_gimp_003
Domain: libreoffice_writer + gimp (multi-app)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_to_gimp_003'

SPECS_FILE = f'{WORKDIR}/edit_specs.docx'
PHOTO_FILE = f'{WORKDIR}/product_photo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_edit_specs():
    """Create the edit_specs.docx with image editing instructions."""
    doc = Document()

    # Title
    title = doc.add_heading('Image Editing Specifications', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        'Please apply the following modifications to the product photo in sequence. '
        'Save the final result as product_photo_edited.png on the Desktop.'
    )
    intro.paragraph_format.space_after = Pt(6)

    # Blank line
    doc.add_paragraph('')

    # Steps heading
    doc.add_heading('Editing Steps', level=2)

    # Step 1
    step1_heading = doc.add_paragraph()
    run = step1_heading.add_run('Step 1: Rotate the image')
    run.bold = True
    run.font.size = Pt(12)

    step1_detail = doc.add_paragraph(
        'Rotate the image 90 degrees clockwise. Use GIMP\'s Image > Transform > '
        'Rotate 90° clockwise function to perform this operation.'
    )
    step1_detail.paragraph_format.left_indent = Pt(18)
    step1_detail.paragraph_format.space_after = Pt(6)

    # Step 2
    step2_heading = doc.add_paragraph()
    run2 = step2_heading.add_run('Step 2: Add a red border')
    run2.bold = True
    run2.font.size = Pt(12)

    step2_detail = doc.add_paragraph(
        'Add a solid red border of 10 pixels around the entire image. '
        'Use GIMP\'s Script-Fu or Filters > Script-Fu > Console to add the border: '
        'flatten the image, then use the canvas size tool to extend the canvas by 20 pixels '
        '(10 on each side) and fill the new border area with red color (255, 0, 0).'
    )
    step2_detail.paragraph_format.left_indent = Pt(18)
    step2_detail.paragraph_format.space_after = Pt(6)

    # Summary section
    doc.add_paragraph('')
    doc.add_heading('Summary of Changes', level=2)

    summary = doc.add_paragraph()
    summary.add_run('Rotation: ').bold = True
    summary.add_run('90 degrees clockwise')
    summary.paragraph_format.space_after = Pt(4)

    summary2 = doc.add_paragraph()
    summary2.add_run('Border: ').bold = True
    summary2.add_run('10-pixel solid red border (RGB: 255, 0, 0)')
    summary2.paragraph_format.space_after = Pt(4)

    summary3 = doc.add_paragraph()
    summary3.add_run('Output file: ').bold = True
    summary3.add_run('product_photo_edited.png (save on the Desktop)')
    summary3.paragraph_format.space_after = Pt(4)

    # Note
    doc.add_paragraph('')
    note = doc.add_paragraph()
    note_run = note.add_run('Note: ')
    note_run.bold = True
    note_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    note.add_run(
        'Apply the steps in the order listed above. Complete Step 1 (rotation) '
        'before Step 2 (border). The final output must include both modifications.'
    )

    doc.save(SPECS_FILE)
    print(f'Specs file created: {SPECS_FILE}')


def create_product_photo():
    """Create a realistic-looking product photo (coffee mug)."""
    width, height = 400, 300

    img = Image.new('RGB', (width, height), color=(240, 235, 225))
    draw = ImageDraw.Draw(img)

    # Background gradient effect (simple)
    for y in range(height):
        shade = int(240 - (y / height) * 30)
        draw.line([(0, y), (width, y)], fill=(shade, shade - 5, shade - 15))

    # Table surface
    draw.rectangle([0, 220, width, height], fill=(180, 140, 100))
    draw.line([(0, 220), (width, 220)], fill=(150, 110, 70), width=2)

    # Mug body
    mug_x1, mug_y1, mug_x2, mug_y2 = 130, 80, 280, 220
    draw.rectangle([mug_x1, mug_y1, mug_x2, mug_y2], fill=(45, 85, 160), outline=(30, 60, 120), width=2)

    # Mug rim (top ellipse)
    draw.ellipse([mug_x1 - 5, mug_y1 - 10, mug_x2 + 5, mug_y1 + 10],
                 fill=(55, 95, 170), outline=(30, 60, 120), width=2)

    # Mug bottom (bottom ellipse)
    draw.ellipse([mug_x1 + 5, mug_y2 - 8, mug_x2 - 5, mug_y2 + 8],
                 fill=(35, 70, 140), outline=(25, 50, 100), width=2)

    # Handle
    handle_x = mug_x2 + 2
    draw.arc([handle_x, 120, handle_x + 50, 200], start=315, end=135,
             fill=(30, 60, 120), width=12)

    # Coffee in mug (dark circle at top)
    draw.ellipse([mug_x1 + 5, mug_y1 - 8, mug_x2 - 5, mug_y1 + 8],
                 fill=(65, 35, 15))

    # Logo / branding on mug: a white star-like design
    cx, cy = (mug_x1 + mug_x2) // 2, (mug_y1 + mug_y2) // 2
    # Circle background for logo
    draw.ellipse([cx - 28, cy - 28, cx + 28, cy + 28],
                 fill=(255, 255, 255), outline=(200, 200, 200), width=1)
    # Star polygon
    import math
    star_points = []
    for i in range(10):
        angle = math.pi / 5 * i - math.pi / 2
        r = 22 if i % 2 == 0 else 10
        star_points.append((cx + r * math.cos(angle), cy + r * math.sin(angle)))
    draw.polygon(star_points, fill=(255, 180, 0), outline=(200, 130, 0))

    # Steam lines above mug
    for offset in [-20, 0, 20]:
        x_base = cx + offset
        for i in range(5):
            y_start = mug_y1 - 15 - i * 8
            y_end = y_start - 6
            x_shift = 3 if i % 2 == 0 else -3
            draw.line([(x_base, y_start), (x_base + x_shift, y_end)],
                      fill=(200, 200, 210), width=2)

    # Product label below
    try:
        from PIL import ImageFont
        font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'
        if os.path.exists(font_path):
            font_large = ImageFont.truetype(font_path, 18)
            font_small = ImageFont.truetype(font_path, 11)
        else:
            font_large = ImageFont.load_default()
            font_small = ImageFont.load_default()
    except Exception:
        font_large = ImageFont.load_default()
        font_small = ImageFont.load_default()

    # Text shadow for product name
    draw.text((149, 231), 'BrewMaster Pro', fill=(120, 90, 50), font=font_large)
    draw.text((148, 230), 'BrewMaster Pro', fill=(255, 255, 255), font=font_large)

    draw.text((162, 254), 'Premium Ceramic', fill=(160, 120, 70), font=font_small)

    img.save(PHOTO_FILE, 'PNG')
    print(f'Product photo created: {PHOTO_FILE}')


def main():
    os.makedirs(WORKDIR, exist_ok=True)

    create_product_photo()
    create_edit_specs()

    # GUI-ready startup: open both the docx and the image
    # First open the docx so the agent can read instructions
    launch_gui(f'libreoffice --writer "{SPECS_FILE}"', delay_sec=2.5)
    # Also open GIMP with the product photo so the agent can edit it
    launch_gui(f'gimp "{PHOTO_FILE}"', delay_sec=2.0)

    print('GUI_READY: launched LibreOffice Writer (edit_specs.docx) and GIMP (product_photo.png) with DISPLAY=:0')


main()
