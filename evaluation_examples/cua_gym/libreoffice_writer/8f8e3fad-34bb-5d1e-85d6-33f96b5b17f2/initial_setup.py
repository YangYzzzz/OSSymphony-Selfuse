"""
Initial Setup: Create Illustrated_Guide.docx with 3 chapters and 12 sequentially numbered figure captions
Task ID: writer_mt_088
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image as PILImage

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_088'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_placeholder_image(path, width=400, height=250, label="Figure"):
    """Create a simple placeholder image with a label."""
    img = PILImage.new('RGB', (width, height), color=(220, 230, 240))
    # Draw a simple border
    pixels = img.load()
    for x in range(width):
        for y in [0, 1, height - 1, height - 2]:
            pixels[x, y] = (100, 120, 150)
    for y in range(height):
        for x in [0, 1, width - 1, width - 2]:
            pixels[x, y] = (100, 120, 150)
    img.save(path)


def create_initial():
    doc = Document()

    # -- Document title --
    title = doc.add_heading('Illustrated Guide to Digital Photography', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This comprehensive guide covers the essential techniques and workflows '
        'for digital photography, from basic camera operation through advanced '
        'post-processing. Each chapter includes detailed illustrations to help '
        'you master the concepts discussed.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Chapter data: (chapter_title, intro_text, figures)
    chapters = [
        (
            'Chapter 1: Introduction to Digital Photography',
            'Digital photography has transformed the way we capture and share moments. '
            'Understanding the fundamentals of camera operation, exposure, and composition '
            'is essential for producing compelling images. This chapter covers the core '
            'concepts every photographer should know.',
            [
                ('Camera sensor diagram showing pixel arrangement and Bayer filter pattern',
                 'The camera sensor converts light into electrical signals through millions of photosites arranged in a grid pattern.'),
                ('Exposure triangle illustrating the relationship between aperture, shutter speed, and ISO',
                 'The exposure triangle demonstrates how the three primary exposure controls work together to determine image brightness.'),
                ('Rule of thirds grid overlay on a landscape photograph',
                 'Composition guidelines such as the rule of thirds help photographers create more visually balanced images.'),
                ('White balance comparison showing the same scene under different color temperature settings',
                 'White balance correction ensures that neutral colors appear accurately regardless of the ambient lighting conditions.'),
            ]
        ),
        (
            'Chapter 2: Advanced Shooting Techniques',
            'Once you have mastered the basics, advanced techniques allow you to tackle '
            'challenging lighting conditions, fast-moving subjects, and creative compositions. '
            'This chapter explores professional methods for elevating your photography.',
            [
                ('Long exposure photograph of a waterfall with silky smooth water effect',
                 'Long exposure techniques create dramatic effects by allowing motion blur to smooth flowing water and clouds.'),
                ('HDR bracketing sequence showing three exposures merged into a single image',
                 'High dynamic range imaging combines multiple exposures to capture detail in both shadows and highlights.'),
                ('Focus stacking result showing front-to-back sharpness in a macro flower photograph',
                 'Focus stacking merges images taken at different focal distances to achieve extended depth of field.'),
            ]
        ),
        (
            'Chapter 3: Post-Processing Workflow',
            'Post-processing is where digital photographs are refined and finalized. '
            'A structured workflow ensures consistent quality and efficient editing. '
            'This chapter walks through the complete post-processing pipeline from '
            'raw file import to final export.',
            [
                ('Raw conversion interface showing histogram and tone curve adjustments',
                 'Raw conversion is the first step in post-processing, allowing non-destructive adjustments to exposure and color.'),
                ('Before and after comparison of color grading applied to a portrait photograph',
                 'Color grading adds mood and atmosphere by shifting color tones in shadows, midtones, and highlights independently.'),
                ('Retouching tools demonstration showing clone stamp and healing brush techniques',
                 'Retouching tools allow photographers to remove distracting elements and refine skin texture in portraits.'),
                ('Sharpening comparison at 100% zoom showing the effect of unsharp mask settings',
                 'Output sharpening compensates for softness introduced during raw conversion and prepares images for display or print.'),
                ('Export dialog showing recommended settings for web and print output formats',
                 'Export settings determine the final file format, resolution, and compression quality for the intended output medium.'),
            ]
        ),
    ]

    # Create placeholder images
    img_paths = []
    for i in range(12):
        path = f'{WORKDIR}/fig_placeholder_{i+1}.png'
        create_placeholder_image(path, label=f"Figure {i+1}")
        img_paths.append(path)

    fig_counter = 0

    for ch_idx, (ch_title, ch_intro, figures) in enumerate(chapters, 1):
        # Chapter heading
        doc.add_heading(ch_title, level=1)

        # Chapter introduction
        doc.add_paragraph(ch_intro)

        for fig_desc, fig_explanation in figures:
            fig_counter += 1

            # Add the figure image
            doc.add_picture(img_paths[fig_counter - 1], width=Inches(4.5))
            # Center the image
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add caption as a separate paragraph with "Caption" style-like formatting
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption_run = caption_para.add_run(f'Figure {fig_counter}: {fig_desc}')
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

            # Explanation paragraph
            doc.add_paragraph(fig_explanation)
            doc.add_paragraph('')  # spacing

    # Final section
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'Digital photography is both a technical discipline and a creative art form. '
        'By understanding the principles covered in this guide and practicing the '
        'techniques illustrated throughout, you will develop the skills needed to '
        'capture and produce professional-quality photographs consistently.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up placeholder images
    for path in img_paths:
        try:
            os.remove(path)
        except OSError:
            pass

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
