"""
Initial Setup: History essay with maritime trade content, no footnotes
Task ID: writer_creative_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_paragraph_format(para, font_name='Times New Roman', font_size=12, line_spacing=2.0):
    """Apply consistent formatting to a paragraph."""
    para.paragraph_format.line_spacing = line_spacing
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def add_formatted_paragraph(doc, text, font_name='Times New Roman', font_size=12, line_spacing=2.0, bold=False):
    """Add a paragraph with specified formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.space_after = Pt(0)
    return para


def create_initial():
    doc = Document()

    # Set default page margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('The Impact of Maritime Trade on Colonial Expansion')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.line_spacing = 2.0
    title_para.paragraph_format.space_after = Pt(0)

    # Paragraph 1 - Introduction
    p1_text = (
        'The relationship between maritime trade and colonial expansion during the early modern period '
        'represents one of the most transformative dynamics in world history. From the fifteenth century '
        'onward, European powers recognized that control of sea lanes translated directly into economic '
        'and political dominance over vast territories. The establishment of trading posts, followed by '
        'permanent settlements, fundamentally reshaped global commerce and cultural exchange in ways that '
        'continue to reverberate into the present day.'
    )
    add_formatted_paragraph(doc, p1_text)

    # Paragraph 2 - contains 'according to Smith (2019)'
    p2_text = (
        'Historians have long debated the precise mechanisms by which commercial interests drove territorial '
        'ambitions. According to recent scholarship, the motivations were rarely purely economic; ideological '
        'and religious factors played equally significant roles. Nevertheless, according to Smith (2019), the '
        'financial returns from spice routes and precious metal extraction provided the material foundation '
        'upon which colonial enterprises were built and sustained over multiple generations. The Portuguese '
        'model of establishing fortified trading posts along the African coast and in the Indian Ocean basin '
        'demonstrated that military superiority at sea could guarantee commercial monopolies of extraordinary '
        'profitability.'
    )
    add_formatted_paragraph(doc, p2_text)

    # Paragraph 3 - contains 'Davis argues'
    p3_text = (
        'The social consequences of this expansion were profound and deeply contested. Colonial powers '
        'disrupted existing trade networks that had functioned for centuries, displacing local merchants '
        'and artisans while creating new dependencies on European manufactured goods. Davis argues that '
        'the indigenous populations of colonized regions were not passive recipients of these changes but '
        'active agents who negotiated, resisted, and occasionally exploited colonial relationships for their '
        'own purposes. This perspective challenges earlier historiography that portrayed colonialism as a '
        'unidirectional imposition of European will upon vulnerable societies lacking the capacity for '
        'meaningful response.'
    )
    add_formatted_paragraph(doc, p3_text)

    # Paragraph 4 - contains 'the Thompson study'
    p4_text = (
        'Quantitative approaches have added considerable precision to our understanding of maritime trade '
        'volumes and their economic impact. Shipbuilding records, customs registers, and merchant account '
        'books preserved in European and Asian archives have enabled scholars to reconstruct trade flows '
        'with increasing accuracy. The Thompson study examined over three thousand voyages undertaken '
        'between 1580 and 1720, revealing patterns of seasonal variation, commodity specialization, and '
        'network resilience that were previously obscured by reliance on qualitative sources alone. '
        'These findings suggest that the scale of intercontinental commerce achieved by the early '
        'seventeenth century exceeded contemporary estimates by a considerable margin.'
    )
    add_formatted_paragraph(doc, p4_text)

    # Paragraph 5 - Conclusion
    p5_text = (
        'In conclusion, the nexus between maritime trade and colonial expansion cannot be understood through '
        'any single explanatory framework. Economic imperatives, technological capabilities, political '
        'rivalries, and cultural assumptions all intersected to produce the distinctive pattern of European '
        'overseas expansion that emerged between 1450 and 1750. Future research must continue to integrate '
        'perspectives from the full range of societies involved in this global transformation, moving beyond '
        'Eurocentric narratives to reconstruct the agency and experience of all participants in this world-'
        'historical process.'
    )
    add_formatted_paragraph(doc, p5_text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
