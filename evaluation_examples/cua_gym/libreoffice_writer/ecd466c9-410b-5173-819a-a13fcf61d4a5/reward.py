"""
FINAL REWARD SCRIPT - SUCCESS
Task: Format the 2 in "Na2CO3" as a subscript in the materials list.
Generated: 2025-10-14 06:49:21
Status: success
Model: azure-o3
Total Steps: 1
"""

from docx import Document
import os

def verify_na2co3_subscript(file_path: str) -> float:
    """Verify that the digit 2 in the chemical formula 'Na2CO3' is
    formatted as a subscript (or supplied as Unicode subscript ₂) somewhere
    in the document (paragraphs or tables).

    Returns a progressive score between 0.0 and 1.0.
    - 1.0  : The digit “2” is correctly formatted as *subscript*.
    - 0.8  : A Unicode subscript digit ₂ is used instead of formatting.
    - 0.0  : Neither formatting nor Unicode subscript detected.
    """

    # ---------- Safety checks ----------
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- Load document ----------
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Failed to open DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Flags for different success levels
    unicode_success = False   # Na₂CO3 (Unicode subscript)
    subscript_success = False # Na2CO3 with run.font.subscript

    # -----------------------------------
    # Helper to process a single paragraph
    # -----------------------------------
    def process_paragraph(paragraph):
        nonlocal unicode_success, subscript_success

        # Build a list pairing each character with its subscript property
        char_info = []  # (char, is_subscript_boolean)
        for run in paragraph.runs:
            is_sub = False
            try:
                # run.font.subscript can be True/False/None
                is_sub = bool(run.font.subscript)
            except Exception:
                is_sub = False
            for ch in run.text:
                char_info.append((ch, is_sub))

        # Skip empty paragraphs
        if not char_info:
            return

        # Full plain text of the paragraph
        text = ''.join(ch for ch, _ in char_info)

        # Look for every occurrence of the literal string 'Na2CO3'
        search_term = 'Na2CO3'
        start_idx = 0
        while True:
            idx = text.find(search_term, start_idx)
            if idx == -1:
                break

            # Examine the third character (the '2') of this occurrence
            ch, is_sub = char_info[idx + 2]
            if ch == '₂':
                unicode_success = True
            elif ch == '2' and is_sub:
                subscript_success = True
            start_idx = idx + 1  # Continue searching after this position

    # ---------- Scan document (paragraphs & tables) ----------
    for para in doc.paragraphs:
        process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    # ---------- Scoring ----------
    if subscript_success:
        score = 1.0
        print("✓ '2' is properly formatted as subscript in Na2CO3 (full credit)")
    elif unicode_success:
        score = 0.8
        print("✓ Unicode subscript ₂ used in Na₂CO3 (partial credit: 0.8)")
    else:
        score = 0.0
        print("✗ No subscript formatting (or Unicode) detected for '2' in Na2CO3")

    print(f"REWARD: {score}")
    return score

# -----------------------------------------------------------
# Execute verification when the script is run directly
# -----------------------------------------------------------
if __name__ == "__main__":
    DOC_PATH = "/home/user/format_the_2_in_na2co3_as_a_subscript_in_the_materials_list.docx"
    verify_na2co3_subscript(DOC_PATH)

