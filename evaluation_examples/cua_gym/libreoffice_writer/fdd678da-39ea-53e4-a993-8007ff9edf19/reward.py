"""
FINAL REWARD SCRIPT - SUCCESS
Task: Please add "Heiss, C., et al. (2003). Improvement of endothelial function with cocoa. Circulation, 107(13), 1652–1658." to the bibliography and reference it by number at the "<add here>" marker in the results section’s opening paragraph.
Generated: 2025-10-14 10:57:57
Status: success
Model: azure-o3
Total Steps: 1
"""

import os
import re
from docx import Document


def verify_task(file_path: str) -> float:
    """Verify that the document fulfills the task requirements.

    Task requirements:
    1. Add the Heiss 2003 reference to the bibliography.
    2. Remove the "<add here>" placeholder that marked where the citation should go.
    3. Insert a numbered (numeric) citation in the opening paragraph of the Results section.

    The function returns a progressive score between 0.0 and 1.0 based on
    successful completion of each requirement:
        • 0.50  – Heiss 2003 reference present somewhere in the document.
        • 0.25  – Placeholder "<add here>" has been removed.
        • 0.25  – A numeric citation (e.g. [3]) is present in the
                  opening paragraph of the Results section.
    """

    print(f"Verifying document: {file_path}")

    # --- Preliminary checks -------------------------------------------------
    if not os.path.exists(file_path):
        print("✗ File not found")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Failed to load DOCX: {exc}")
        return 0.0  # Cannot continue without loading

    # Gather all text in a single string for global searches
    all_text = "\n".join(p.text for p in doc.paragraphs if p.text)
    paragraphs = list(doc.paragraphs)

    score = 0.0

    # --- Requirement 1: Heiss 2003 reference present ------------------------
    citation_regex = re.compile(
        r"Heiss\s*,\s*C\.?,?\s*et\s*al\.?.*?2003.*?endothelial\s+function.*?cocoa",
        re.IGNORECASE | re.DOTALL,
    )

    if citation_regex.search(all_text):
        print("✓ Heiss 2003 citation found in document")
        score += 0.5
    else:
        print("✗ Heiss 2003 citation NOT found in document")

    # --- Requirement 2: Placeholder removed ---------------------------------
    if re.search(r"<\s*add\s+here\s*>", all_text, re.IGNORECASE):
        print("✗ Placeholder '<add here>' still present in document")
    else:
        print("✓ Placeholder '<add here>' successfully removed")
        score += 0.25

    # --- Requirement 3: Numeric citation in Results opening paragraph -------
    # Locate the 'Results' heading (with or without numbering like '3. Results')
    results_heading_idx = None
    for idx, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        if text == "results" or re.fullmatch(r"\d+\.\s*results", text):
            results_heading_idx = idx
            break

    numeric_citation_found = False
    if results_heading_idx is not None:
        # Find first non-empty paragraph after the heading
        for j in range(results_heading_idx + 1, len(paragraphs)):
            opening_para_text = paragraphs[j].text.strip()
            if opening_para_text:
                print("Results opening paragraph (truncated):",
                      opening_para_text[:120])
                # Patterns for numeric citation: [n], (n) or superscript digits
                if re.search(r"\[\s*\d+\s*\]", opening_para_text):
                    numeric_citation_found = True
                elif re.search(r"\(\s*\d+\s*\)", opening_para_text):
                    numeric_citation_found = True
                else:
                    # Check for superscript digits (¹²³⁴⁵⁶⁷⁸⁹⁰)
                    superscripts = "¹²³⁴⁵⁶⁷⁸⁹⁰"
                    if any(ch in opening_para_text for ch in superscripts):
                        numeric_citation_found = True
                break  # Only evaluate the first non-empty paragraph
    else:
        print("✗ 'Results' heading not found – cannot verify numeric citation")

    if numeric_citation_found:
        print("✓ Numeric citation present in Results opening paragraph")
        score += 0.25
    else:
        print("✗ Numeric citation NOT found in Results opening paragraph")

    # --- Final score --------------------------------------------------------
    final_score = min(score, 1.0)
    print(f"Total score: {final_score}")
    return final_score


if __name__ == "__main__":
    DOC_PATH = (
        "/home/user/"
        "please_add_heiss_c_et_al_2003_improvement_of_endothelial_"
        "function_with_cocoa_circulation_10713_16521.docx"
    )
    reward_value = verify_task(DOC_PATH)
    print(f"REWARD: {reward_value}")
