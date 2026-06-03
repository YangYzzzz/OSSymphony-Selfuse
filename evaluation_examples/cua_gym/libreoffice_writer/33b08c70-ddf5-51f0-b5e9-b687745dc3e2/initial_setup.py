"""
Initial Setup: Create a contract document in single-column layout with a Definitions section
Task ID: writer_legal_054
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_054'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---- Title ----
    title = doc.add_heading('SOFTWARE LICENSING AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Preamble ----
    doc.add_heading('1. PREAMBLE', level=1)
    doc.add_paragraph(
        'This Software Licensing Agreement ("Agreement") is entered into as of March 15, 2025, '
        'by and between Meridian Technologies Inc., a Delaware corporation with its principal '
        'offices at 2400 Innovation Boulevard, Suite 800, San Francisco, CA 94105 ("Licensor"), '
        'and Cascade Digital Solutions LLC, a Washington limited liability company with its '
        'principal offices at 1750 Pacific Avenue, Seattle, WA 98101 ("Licensee").'
    )
    doc.add_paragraph(
        'WHEREAS, Licensor has developed certain proprietary software known as MeridianFlow '
        'Enterprise Platform (the "Software"); and WHEREAS, Licensee desires to obtain a license '
        'to use the Software for its internal business operations subject to the terms and '
        'conditions set forth herein.'
    )
    doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, '
        'and for other good and valuable consideration, the receipt and sufficiency of which are '
        'hereby acknowledged, the parties agree as follows:'
    )

    # ---- Definitions ----
    doc.add_heading('2. DEFINITIONS', level=1)

    definitions = [
        ('"Affiliate"', 'means any entity that directly or indirectly controls, is controlled by, '
         'or is under common control with a party, where "control" means ownership of more than '
         'fifty percent (50%) of the voting securities or equivalent ownership interest.'),
        ('"Authorized Users"', 'means the employees, contractors, and agents of Licensee who are '
         'authorized by Licensee to access and use the Software under this Agreement, subject to '
         'the user limitations specified in the applicable Order Form.'),
        ('"Business Day"', 'means any day other than a Saturday, Sunday, or public holiday in the '
         'State of California or the State of Washington.'),
        ('"Confidential Information"', 'means all non-public information disclosed by either party '
         'to the other party, whether orally, in writing, or by inspection, including but not '
         'limited to trade secrets, business plans, financial data, and technical specifications.'),
        ('"Data Processing Agreement"', 'means the data processing addendum attached hereto as '
         'Exhibit B, governing the processing of personal data by Licensor on behalf of Licensee.'),
        ('"Documentation"', 'means the user manuals, online help files, technical specifications, '
         'and other written materials provided by Licensor to describe the functionality and '
         'operation of the Software.'),
        ('"Effective Date"', 'means March 15, 2025, or such later date as specified in the '
         'applicable Order Form executed by both parties.'),
        ('"Force Majeure Event"', 'means any event beyond the reasonable control of the affected '
         'party, including but not limited to acts of God, natural disasters, epidemics, war, '
         'terrorism, government actions, or failures of third-party telecommunications networks.'),
        ('"Intellectual Property Rights"', 'means all patents, copyrights, trademarks, trade '
         'secrets, and other proprietary rights recognized under applicable law, whether registered '
         'or unregistered, and all applications and registrations therefor.'),
        ('"License Fee"', 'means the fees payable by Licensee to Licensor for the rights granted '
         'under this Agreement, as specified in the Order Form, currently set at $245,000 per annum.'),
        ('"Maintenance Release"', 'means a release of the Software that corrects errors, bugs, or '
         'defects, and is designated by Licensor as a maintenance or patch release (e.g., version '
         'changes in the third digit, such as from 4.2.1 to 4.2.2).'),
        ('"Order Form"', 'means the ordering document executed by both parties that references this '
         'Agreement and specifies the Software modules, number of Authorized Users, License Fees, '
         'and subscription term.'),
        ('"Personal Data"', 'means any information relating to an identified or identifiable natural '
         'person as defined under applicable data protection legislation, including the California '
         'Consumer Privacy Act (CCPA) and the General Data Protection Regulation (GDPR).'),
        ('"Professional Services"', 'means implementation, configuration, customization, training, '
         'and consulting services provided by Licensor to Licensee as described in the applicable '
         'Statement of Work.'),
        ('"SLA"', 'means the Service Level Agreement attached hereto as Exhibit A, specifying '
         'uptime commitments of 99.9% availability, response times, and remedies for service '
         'level failures.'),
        ('"Software"', 'means the MeridianFlow Enterprise Platform, including all modules, '
         'components, updates, upgrades, and Maintenance Releases provided by Licensor during '
         'the term of this Agreement.'),
        ('"Statement of Work"', 'means a document executed by both parties describing the scope, '
         'timeline, deliverables, and fees for Professional Services to be performed by Licensor.'),
        ('"Subscription Term"', 'means the period during which Licensee is authorized to use the '
         'Software, as specified in the applicable Order Form, initially set at thirty-six (36) '
         'months from the Effective Date.'),
        ('"Third-Party Components"', 'means any software, libraries, or modules owned by third '
         'parties that are incorporated into or distributed with the Software, as listed in '
         'Exhibit C attached hereto.'),
        ('"Update"', 'means a release of the Software that includes new features, functionality '
         'enhancements, or significant improvements, and is designated by Licensor as a major or '
         'minor version release (e.g., version changes from 4.2 to 4.3 or from 4.x to 5.0).'),
    ]

    for term, definition in definitions:
        para = doc.add_paragraph()
        run_term = para.add_run(term)
        run_term.bold = True
        para.add_run(f' {definition}')

    # ---- Grant of License ----
    doc.add_heading('3. GRANT OF LICENSE', level=1)
    doc.add_paragraph(
        'Subject to the terms and conditions of this Agreement and timely payment of all License '
        'Fees, Licensor hereby grants to Licensee a non-exclusive, non-transferable, limited '
        'license to install, access, and use the Software during the Subscription Term solely '
        'for Licensee\'s internal business operations.'
    )
    doc.add_paragraph(
        'The license granted herein is limited to the number of Authorized Users specified in '
        'the Order Form. Licensee shall not permit access to the Software by any person other '
        'than Authorized Users without the prior written consent of Licensor.'
    )

    # ---- Fees and Payment ----
    doc.add_heading('4. FEES AND PAYMENT', level=1)
    doc.add_paragraph(
        'Licensee shall pay the License Fees as specified in the Order Form within thirty (30) '
        'days of invoice date. All fees are quoted in United States Dollars and are non-refundable '
        'except as expressly provided in this Agreement. Late payments shall accrue interest at the '
        'rate of 1.5% per month or the maximum rate permitted by applicable law, whichever is lower.'
    )
    doc.add_paragraph(
        'Licensor reserves the right to increase License Fees upon renewal of the Subscription Term '
        'by providing written notice at least ninety (90) days prior to the renewal date. Any fee '
        'increase shall not exceed eight percent (8%) of the then-current License Fees.'
    )

    # ---- Limitation of Liability ----
    doc.add_heading('5. LIMITATION OF LIABILITY', level=1)
    doc.add_paragraph(
        'IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER PARTY FOR ANY INDIRECT, INCIDENTAL, '
        'SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING BUT NOT LIMITED TO LOSS OF PROFITS, '
        'LOSS OF DATA, BUSINESS INTERRUPTION, OR LOSS OF GOODWILL, ARISING OUT OF OR IN CONNECTION '
        'WITH THIS AGREEMENT, REGARDLESS OF THE THEORY OF LIABILITY.'
    )
    doc.add_paragraph(
        'The aggregate liability of Licensor under this Agreement shall not exceed the total amount '
        'of License Fees actually paid by Licensee during the twelve (12) month period immediately '
        'preceding the event giving rise to such liability.'
    )

    # ---- Termination ----
    doc.add_heading('6. TERMINATION', level=1)
    doc.add_paragraph(
        'Either party may terminate this Agreement upon sixty (60) days\' written notice to the '
        'other party if the other party commits a material breach of this Agreement and fails to '
        'cure such breach within thirty (30) days after receiving written notice thereof.'
    )
    doc.add_paragraph(
        'Upon termination or expiration of this Agreement, Licensee shall immediately cease all '
        'use of the Software, destroy all copies of the Software and Documentation in its '
        'possession, and certify such destruction in writing to Licensor within fifteen (15) '
        'Business Days.'
    )

    # ---- Governing Law ----
    doc.add_heading('7. GOVERNING LAW', level=1)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of the '
        'State of California, without regard to its conflict of laws principles. Any dispute '
        'arising out of or relating to this Agreement shall be resolved exclusively in the state '
        'or federal courts located in San Francisco County, California.'
    )

    # ---- Signatures ----
    doc.add_heading('8. SIGNATURES', level=1)
    doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.'
    )

    doc.add_paragraph('')
    sig1 = doc.add_paragraph()
    sig1.add_run('For Meridian Technologies Inc.:').bold = True
    doc.add_paragraph('Name: Victoria R. Harrington')
    doc.add_paragraph('Title: Chief Executive Officer')
    doc.add_paragraph('Date: March 15, 2025')

    doc.add_paragraph('')
    sig2 = doc.add_paragraph()
    sig2.add_run('For Cascade Digital Solutions LLC:').bold = True
    doc.add_paragraph('Name: Daniel K. Otsuka')
    doc.add_paragraph('Title: Managing Director')
    doc.add_paragraph('Date: March 15, 2025')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
