"""
Initial Setup: Research paper document with empty bibliography database
Task ID: writer_acad_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('Advances in Neural Network Architectures for Natural Language Processing', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Authors ---
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Elena Rodriguez, Wei Zhang, and Priya Nair')
    run.font.size = Pt(12)
    run.font.italic = True

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = affil.add_run('Department of Computer Science, Stanford University')
    run2.font.size = Pt(10)

    # --- Abstract ---
    doc.add_heading('Abstract', level=2)
    abstract_para = doc.add_paragraph()
    abstract_para.paragraph_format.first_line_indent = Inches(0.25)
    abstract_text = (
        'This paper presents a comprehensive survey of recent advances in neural network '
        'architectures designed for natural language processing tasks. We examine transformer-based '
        'models, attention mechanisms, and their applications across sentiment analysis, machine '
        'translation, and text summarization. Our analysis covers over 150 publications from 2018 '
        'to 2024 and identifies key trends in model efficiency, scalability, and multilingual '
        'capabilities. We find that while transformer architectures continue to dominate, emerging '
        'hybrid approaches combining convolutional and recurrent elements show promising results '
        'for resource-constrained environments.'
    )
    run_abs = abstract_para.add_run(abstract_text)
    run_abs.font.size = Pt(11)

    # --- Keywords ---
    kw_para = doc.add_paragraph()
    run_kw_label = kw_para.add_run('Keywords: ')
    run_kw_label.bold = True
    run_kw_label.font.size = Pt(11)
    run_kw = kw_para.add_run(
        'natural language processing, transformer models, attention mechanisms, deep learning, '
        'neural architectures'
    )
    run_kw.font.size = Pt(11)
    run_kw.font.italic = True

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=2)

    intro1 = doc.add_paragraph()
    intro1.paragraph_format.first_line_indent = Inches(0.25)
    intro1.paragraph_format.space_after = Pt(6)
    r = intro1.add_run(
        'Natural language processing (NLP) has undergone a remarkable transformation in recent '
        'years, driven primarily by advances in deep learning and the introduction of transformer '
        'architectures. The publication of the original transformer model by Vaswani et al. in '
        '2017 marked a paradigm shift, moving the field away from recurrent neural networks '
        'toward attention-based approaches that enable significantly greater parallelization '
        'during training.'
    )
    r.font.size = Pt(11)

    intro2 = doc.add_paragraph()
    intro2.paragraph_format.first_line_indent = Inches(0.25)
    intro2.paragraph_format.space_after = Pt(6)
    r = intro2.add_run(
        'Subsequent developments, including BERT, GPT, and their numerous variants, have '
        'achieved state-of-the-art performance across virtually all NLP benchmarks. These models '
        'leverage large-scale pre-training on diverse corpora followed by task-specific '
        'fine-tuning, a paradigm that has proven remarkably effective for tasks ranging from '
        'question answering to text generation and document classification.'
    )
    r.font.size = Pt(11)

    intro3 = doc.add_paragraph()
    intro3.paragraph_format.first_line_indent = Inches(0.25)
    intro3.paragraph_format.space_after = Pt(6)
    r = intro3.add_run(
        'Despite these advances, significant challenges remain. Model sizes have grown '
        'exponentially, raising concerns about computational costs, energy consumption, and '
        'accessibility. Furthermore, questions about robustness, fairness, and interpretability '
        'continue to motivate research into alternative architectural designs and training '
        'methodologies.'
    )
    r.font.size = Pt(11)

    # --- 2. Background ---
    doc.add_heading('2. Background and Related Work', level=2)

    bg1 = doc.add_paragraph()
    bg1.paragraph_format.first_line_indent = Inches(0.25)
    bg1.paragraph_format.space_after = Pt(6)
    r = bg1.add_run(
        'The evolution of NLP architectures can be broadly categorized into three eras: '
        'rule-based systems (1950s-1990s), statistical methods (1990s-2010s), and neural '
        'approaches (2010s-present). Early neural models for NLP relied heavily on recurrent '
        'neural networks, particularly Long Short-Term Memory (LSTM) networks introduced by '
        'Hochreiter and Schmidhuber in 1997.'
    )
    r.font.size = Pt(11)

    bg2 = doc.add_paragraph()
    bg2.paragraph_format.first_line_indent = Inches(0.25)
    bg2.paragraph_format.space_after = Pt(6)
    r = bg2.add_run(
        'The attention mechanism, first proposed for neural machine translation by Bahdanau '
        'et al. in 2014, allowed models to selectively focus on relevant parts of the input '
        'sequence. This innovation laid the groundwork for the transformer architecture, which '
        'relies entirely on self-attention and dispenses with recurrence altogether.'
    )
    r.font.size = Pt(11)

    # --- 3. Methodology ---
    doc.add_heading('3. Methodology', level=2)

    meth1 = doc.add_paragraph()
    meth1.paragraph_format.first_line_indent = Inches(0.25)
    meth1.paragraph_format.space_after = Pt(6)
    r = meth1.add_run(
        'Our survey methodology follows a systematic approach. We collected publications '
        'from major venues including ACL, EMNLP, NeurIPS, ICML, and ICLR spanning the years '
        '2018 through 2024. We applied inclusion criteria requiring that papers propose or '
        'substantially modify a neural architecture for at least one NLP task.'
    )
    r.font.size = Pt(11)

    # --- Table: Survey Statistics ---
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Year', 'Papers Reviewed', 'Key Architectures']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['2021', '38', 'DeBERTa, ERNIE 3.0'],
        ['2022', '42', 'PaLM, Chinchilla, OPT'],
        ['2023', '35', 'LLaMA, Mistral, Falcon'],
        ['2024', '28', 'Mamba, RWKV, Jamba'],
    ]
    for r_idx, row_data in enumerate(data, 1):
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    # --- 4. Discussion ---
    doc.add_heading('4. Discussion', level=2)

    disc1 = doc.add_paragraph()
    disc1.paragraph_format.first_line_indent = Inches(0.25)
    disc1.paragraph_format.space_after = Pt(6)
    r = disc1.add_run(
        'Our analysis reveals several notable trends. First, the dominance of transformer-based '
        'architectures continues, with over 85% of surveyed papers either proposing transformer '
        'variants or building upon existing transformer foundations. Second, there is a growing '
        'emphasis on efficiency, with techniques such as sparse attention, mixture of experts, '
        'and knowledge distillation receiving increasing attention.'
    )
    r.font.size = Pt(11)

    disc2 = doc.add_paragraph()
    disc2.paragraph_format.first_line_indent = Inches(0.25)
    disc2.paragraph_format.space_after = Pt(6)
    r = disc2.add_run(
        'Third, state-space models such as Mamba represent a potential paradigm shift, offering '
        'linear-time sequence processing without sacrificing quality on language modeling '
        'benchmarks. Whether these models will ultimately supplant transformers remains an '
        'open question that warrants continued investigation.'
    )
    r.font.size = Pt(11)

    # --- 5. Conclusion ---
    doc.add_heading('5. Conclusion', level=2)

    concl = doc.add_paragraph()
    concl.paragraph_format.first_line_indent = Inches(0.25)
    concl.paragraph_format.space_after = Pt(6)
    r = concl.add_run(
        'This survey has provided a comprehensive overview of neural network architectures '
        'for NLP from 2018 to 2024. The field continues to evolve rapidly, with new '
        'architectures and training strategies emerging at an accelerating pace. Future work '
        'should focus on improving model efficiency, enhancing multilingual capabilities, '
        'and developing more interpretable architectures that can be deployed responsibly '
        'in real-world applications.'
    )
    r.font.size = Pt(11)

    # --- References (placeholder section, no bibliography entries) ---
    doc.add_heading('References', level=2)
    ref_note = doc.add_paragraph()
    r = ref_note.add_run('[Bibliography entries to be added via bibliography database]')
    r.font.size = Pt(10)
    r.font.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
