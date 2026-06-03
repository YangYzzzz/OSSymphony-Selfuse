"""
Initial Setup: Create a finalized legal contract document for PDF export task.
Task ID: writer_legal_066
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_066'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page Setup --
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Title --
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- Contract Number and Date --
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Contract No. PSA-2025-0482')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    meta.add_run('\n')
    run2 = meta.add_run('Effective Date: March 15, 2025')
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph('')  # spacer

    # -- Preamble --
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(6)
    run = preamble.add_run(
        'This Professional Services Agreement ("Agreement") is entered into as of '
        'March 15, 2025 ("Effective Date"), by and between:'
    )
    run.font.size = Pt(11)

    # -- Party A --
    h1 = doc.add_heading('PARTY A (Client)', level=2)
    party_a = doc.add_paragraph()
    party_a.paragraph_format.left_indent = Inches(0.5)
    lines_a = [
        'Meridian Technologies, Inc.',
        '1200 Innovation Drive, Suite 450',
        'San Francisco, CA 94105',
        'Tax ID: 94-3827156',
        'Represented by: Victoria Chen, Chief Technology Officer'
    ]
    for i, line in enumerate(lines_a):
        r = party_a.add_run(line)
        r.font.size = Pt(11)
        if i == 0:
            r.bold = True
        if i < len(lines_a) - 1:
            party_a.add_run('\n')

    # -- Party B --
    h2 = doc.add_heading('PARTY B (Service Provider)', level=2)
    party_b = doc.add_paragraph()
    party_b.paragraph_format.left_indent = Inches(0.5)
    lines_b = [
        'Pinnacle Consulting Group, LLC',
        '875 Market Street, Floor 12',
        'New York, NY 10013',
        'Tax ID: 13-5924817',
        'Represented by: James R. Patterson, Managing Partner'
    ]
    for i, line in enumerate(lines_b):
        r = party_b.add_run(line)
        r.font.size = Pt(11)
        if i == 0:
            r.bold = True
        if i < len(lines_b) - 1:
            party_b.add_run('\n')

    doc.add_paragraph('')  # spacer

    # -- Section 1: Scope of Services --
    doc.add_heading('1. SCOPE OF SERVICES', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        '1.1 The Service Provider agrees to provide the following professional services '
        'to the Client ("Services"):'
    )
    r.font.size = Pt(11)

    services = [
        'Enterprise cloud infrastructure migration and optimization for all North American data centers',
        'Development and deployment of custom API integration layer connecting legacy ERP systems with modern SaaS platforms',
        'Implementation of zero-trust security architecture across all production environments',
        'Staff training program covering DevOps practices, CI/CD pipelines, and incident response protocols',
        'Post-migration performance monitoring and optimization for a period of twelve (12) months'
    ]
    for svc in services:
        bp = doc.add_paragraph(svc, style='List Bullet')
        for run in bp.runs:
            run.font.size = Pt(11)

    p12 = doc.add_paragraph()
    p12.paragraph_format.space_after = Pt(6)
    r = p12.add_run(
        '1.2 The Service Provider shall assign a dedicated project team consisting of '
        'no fewer than eight (8) qualified professionals, including a Senior Project Manager, '
        'two (2) Solutions Architects, three (3) Senior Engineers, and two (2) QA Specialists.'
    )
    r.font.size = Pt(11)

    # -- Section 2: Compensation --
    doc.add_heading('2. COMPENSATION AND PAYMENT TERMS', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        '2.1 The Client agrees to pay the Service Provider a total fixed fee of '
        'Seven Hundred Fifty Thousand United States Dollars ($750,000.00) for the '
        'complete scope of Services described in Section 1.'
    )
    r.font.size = Pt(11)

    # Payment schedule table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Milestone', 'Payment Amount', 'Due Date']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    payments = [
        ['Project Kickoff', '$150,000.00', 'April 1, 2025'],
        ['Phase 1 Completion - Infrastructure Assessment', '$187,500.00', 'June 15, 2025'],
        ['Phase 2 Completion - Migration Execution', '$225,000.00', 'September 30, 2025'],
        ['Final Delivery & Acceptance', '$187,500.00', 'December 15, 2025'],
    ]
    for r_idx, row_data in enumerate(payments, 1):
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph('')

    p22 = doc.add_paragraph()
    p22.paragraph_format.space_after = Pt(6)
    r = p22.add_run(
        '2.2 All invoices shall be submitted via electronic means and are payable within '
        'thirty (30) calendar days of receipt. Late payments shall accrue interest at a '
        'rate of 1.5% per month on the outstanding balance.'
    )
    r.font.size = Pt(11)

    # -- Section 3: Term and Termination --
    doc.add_heading('3. TERM AND TERMINATION', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        '3.1 This Agreement shall commence on the Effective Date and continue for a '
        'period of twelve (12) months, unless terminated earlier in accordance with '
        'the provisions herein.'
    )
    r.font.size = Pt(11)

    p32 = doc.add_paragraph()
    p32.paragraph_format.space_after = Pt(6)
    r = p32.add_run(
        '3.2 Either party may terminate this Agreement with sixty (60) days\' written '
        'notice to the other party. In the event of early termination by the Client, '
        'the Client shall pay for all Services rendered through the termination date '
        'plus a termination fee equal to ten percent (10%) of the remaining contract value.'
    )
    r.font.size = Pt(11)

    # -- Section 4: Confidentiality --
    doc.add_heading('4. CONFIDENTIALITY AND DATA PROTECTION', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        '4.1 Both parties acknowledge that during the performance of this Agreement, '
        'each party may have access to Confidential Information of the other party. '
        '"Confidential Information" includes, but is not limited to, trade secrets, '
        'business strategies, customer lists, financial data, proprietary software code, '
        'technical specifications, and employee information.'
    )
    r.font.size = Pt(11)

    p42 = doc.add_paragraph()
    p42.paragraph_format.space_after = Pt(6)
    r = p42.add_run(
        '4.2 The Service Provider shall comply with all applicable data protection '
        'regulations, including but not limited to the California Consumer Privacy Act '
        '(CCPA), the General Data Protection Regulation (GDPR) where applicable, and '
        'any industry-specific compliance requirements (SOC 2 Type II, ISO 27001).'
    )
    r.font.size = Pt(11)

    # -- Section 5: Intellectual Property --
    doc.add_heading('5. INTELLECTUAL PROPERTY RIGHTS', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        '5.1 All deliverables, work product, documentation, and custom code developed '
        'by the Service Provider specifically for the Client under this Agreement '
        '("Work Product") shall be the exclusive property of the Client upon full payment.'
    )
    r.font.size = Pt(11)

    p52 = doc.add_paragraph()
    p52.paragraph_format.space_after = Pt(6)
    r = p52.add_run(
        '5.2 The Service Provider retains ownership of all pre-existing intellectual '
        'property, tools, frameworks, and methodologies ("Provider IP"). The Service '
        'Provider hereby grants the Client a non-exclusive, perpetual, royalty-free '
        'license to use any Provider IP incorporated into the Work Product.'
    )
    r.font.size = Pt(11)

    # -- Section 6: Liability --
    doc.add_heading('6. LIMITATION OF LIABILITY', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        '6.1 EXCEPT FOR BREACHES OF CONFIDENTIALITY OBLIGATIONS OR INTELLECTUAL '
        'PROPERTY INFRINGEMENT, NEITHER PARTY SHALL BE LIABLE FOR ANY INDIRECT, '
        'INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF '
        'OR RELATING TO THIS AGREEMENT.'
    )
    r.font.size = Pt(11)

    p62 = doc.add_paragraph()
    p62.paragraph_format.space_after = Pt(6)
    r = p62.add_run(
        '6.2 The Service Provider\'s total aggregate liability under this Agreement '
        'shall not exceed the total fees paid by the Client to the Service Provider '
        'during the twelve (12) month period preceding the claim.'
    )
    r.font.size = Pt(11)

    # -- Section 7: Governing Law --
    doc.add_heading('7. GOVERNING LAW AND DISPUTE RESOLUTION', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        '7.1 This Agreement shall be governed by and construed in accordance with the '
        'laws of the State of California, without regard to its conflict of laws principles.'
    )
    r.font.size = Pt(11)

    p72 = doc.add_paragraph()
    p72.paragraph_format.space_after = Pt(6)
    r = p72.add_run(
        '7.2 Any disputes arising under this Agreement shall first be submitted to '
        'mediation in San Francisco, California. If mediation fails to resolve the '
        'dispute within sixty (60) days, either party may initiate binding arbitration '
        'under the rules of the American Arbitration Association.'
    )
    r.font.size = Pt(11)

    # -- Signature Block --
    doc.add_paragraph('')
    doc.add_heading('SIGNATURES', level=1)

    sig_intro = doc.add_paragraph()
    r = sig_intro.add_run(
        'IN WITNESS WHEREOF, the parties have executed this Agreement as of the '
        'Effective Date first written above.'
    )
    r.font.size = Pt(11)
    r.italic = True

    doc.add_paragraph('')

    # Party A signature
    sig_a = doc.add_paragraph()
    sig_a.add_run('_' * 40 + '\n').font.size = Pt(11)
    r = sig_a.add_run('Victoria Chen\n')
    r.font.size = Pt(11)
    r.bold = True
    r2 = sig_a.add_run('Chief Technology Officer\nMeridian Technologies, Inc.\nDate: March 15, 2025')
    r2.font.size = Pt(11)

    doc.add_paragraph('')

    # Party B signature
    sig_b = doc.add_paragraph()
    sig_b.add_run('_' * 40 + '\n').font.size = Pt(11)
    r = sig_b.add_run('James R. Patterson\n')
    r.font.size = Pt(11)
    r.bold = True
    r2 = sig_b.add_run('Managing Partner\nPinnacle Consulting Group, LLC\nDate: March 15, 2025')
    r2.font.size = Pt(11)

    # -- Footer --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = fp.add_run('CONFIDENTIAL - Meridian Technologies, Inc. | PSA-2025-0482')
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
