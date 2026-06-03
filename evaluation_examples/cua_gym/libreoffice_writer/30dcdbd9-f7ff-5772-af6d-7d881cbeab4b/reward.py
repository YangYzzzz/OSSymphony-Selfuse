"""
Reward Script: Insert a rounded rectangle shape behind title with gradient fill and white title text
Task ID: writer_rd_036
Domain: libreoffice_writer
Scoring:
  Component 1: Drawing/shape exists in the document (0.15 pts)
  Component 2: Shape is a rounded rectangle (roundRect preset geometry) (0.15 pts)
  Component 3: Shape is behind text (behindDoc=1) (0.15 pts)
  Component 4: Shape has gradient fill from dark blue #003366 to light blue #6699CC (0.30 pts)
  Component 5: Title text color is white (#FFFFFF) (0.25 pts)
"""

import os
from math import sqrt
from docx import Document
from docx.shared import RGBColor
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_036'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def color_distance(c1, c2):
    """Euclidean distance in RGB space."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def hex_to_rgb(hex_str):
    """Convert hex color string to (R, G, B) tuple."""
    hex_str = hex_str.lstrip('#')
    return (int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    body = doc.element.body
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    }

    # Find all drawings in the document
    drawings = body.findall('.//w:drawing', ns)

    # Component 1: At least one drawing/shape exists (0.15 pts)
    try:
        if len(drawings) > 0:
            print(f"PASS: Component 1 — Found {len(drawings)} drawing(s) in document (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — No drawings found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Find the specific shape elements for deeper checks
    shape_found = False
    is_round_rect = False
    is_behind_doc = False
    has_gradient = False
    gradient_colors_ok = False

    for drawing in drawings:
        # Check for anchor element
        anchors = drawing.findall('.//wp:anchor', ns)
        for anchor in anchors:
            behind_doc = anchor.get('behindDoc', '0')

            # Check preset geometry
            prst_geoms = anchor.findall('.//a:prstGeom', ns)
            for geom in prst_geoms:
                prst = geom.get('prst', '')
                if prst == 'roundRect':
                    is_round_rect = True
                    shape_found = True

            if behind_doc == '1':
                is_behind_doc = True

            # Check gradient fill
            grad_fills = anchor.findall('.//a:gradFill', ns)
            for gf in grad_fills:
                has_gradient = True
                gs_list = gf.findall('.//a:gs', ns)
                colors_found = {}
                for gs in gs_list:
                    pos = gs.get('pos', '')
                    srgb = gs.find('.//a:srgbClr', ns)
                    if srgb is not None:
                        color_val = srgb.get('val', '')
                        colors_found[pos] = color_val

                # Check gradient colors: dark blue at start, light blue at end
                # Allow tolerance for similar positions and colors
                dark_blue_target = hex_to_rgb('003366')
                light_blue_target = hex_to_rgb('6699CC')

                has_dark = False
                has_light = False
                for pos, color_hex in colors_found.items():
                    color_rgb = hex_to_rgb(color_hex)
                    # Check if this is the dark blue (at start position, roughly 0-30%)
                    if int(pos) <= 30000 and color_distance(color_rgb, dark_blue_target) < 50:
                        has_dark = True
                    # Check if this is the light blue (at end position, roughly 70-100%)
                    if int(pos) >= 70000 and color_distance(color_rgb, light_blue_target) < 50:
                        has_light = True

                if has_dark and has_light:
                    gradient_colors_ok = True

        # Also check inline shapes (wp:inline) though less likely for behind-text
        inlines = drawing.findall('.//wp:inline', ns)
        for inline in inlines:
            prst_geoms = inline.findall('.//a:prstGeom', ns)
            for geom in prst_geoms:
                prst = geom.get('prst', '')
                if prst == 'roundRect':
                    is_round_rect = True
                    shape_found = True

    # Component 2: Shape is a rounded rectangle (0.15 pts)
    try:
        if is_round_rect:
            print(f"PASS: Component 2 — Shape has roundRect preset geometry (0.15 pts)")
            total_score += 0.15
        else:
            # Also accept if any shape/rectangle exists (partial tolerance)
            any_shape = len(drawings) > 0
            if any_shape:
                prsts = body.findall('.//a:prstGeom', ns)
                found_prsts = [p.get('prst', '') for p in prsts]
                print(f"FAIL: Component 2 — Expected roundRect, found preset geometries: {found_prsts}")
            else:
                print(f"FAIL: Component 2 — No shape found to check geometry")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Shape is behind text (behindDoc=1) (0.15 pts)
    try:
        if is_behind_doc and len(drawings) > 0:
            print(f"PASS: Component 3 — Shape anchor has behindDoc=1 (behind text) (0.15 pts)")
            total_score += 0.15
        else:
            if len(drawings) > 0:
                # Check wrap mode alternatives
                wrap_types = []
                for drawing in drawings:
                    for anchor in drawing.findall('.//wp:anchor', ns):
                        bd = anchor.get('behindDoc', 'not set')
                        wrap_types.append(f"behindDoc={bd}")
                print(f"FAIL: Component 3 — Shape not behind text. Found: {wrap_types}")
            else:
                print(f"FAIL: Component 3 — No drawings to check wrap mode")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Gradient fill from dark blue to light blue (0.30 pts)
    try:
        if has_gradient and gradient_colors_ok:
            print(f"PASS: Component 4 — Gradient fill with correct colors #003366 → #6699CC (0.30 pts)")
            total_score += 0.30
        elif has_gradient:
            # Partial credit: gradient exists but colors may be off
            print(f"PARTIAL: Component 4 — Gradient fill exists but colors don't match target (0.15 pts)")
            total_score += 0.15
        else:
            # Check if there's any fill at all
            solid_fills = body.findall('.//wps:spPr/a:solidFill', ns)
            print(f"FAIL: Component 4 — No gradient fill found on shape. Solid fills: {len(solid_fills)}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Title text color is white (#FFFFFF) (0.25 pts)
    try:
        if len(doc.paragraphs) > 0:
            title_para = doc.paragraphs[0]
            white_target = (255, 255, 255)
            title_is_white = False

            for run in title_para.runs:
                if run.text.strip():
                    rgb = run.font.color.rgb
                    if rgb is not None:
                        run_color = (rgb[0], rgb[1], rgb[2])
                        dist = color_distance(run_color, white_target)
                        if dist < 30:  # tolerance for near-white
                            title_is_white = True
                            break

            if title_is_white:
                print(f"PASS: Component 5 — Title text color is white/near-white (0.25 pts)")
                total_score += 0.25
            else:
                # Report what the actual color is
                colors = []
                for run in title_para.runs:
                    if run.text.strip():
                        rgb = run.font.color.rgb
                        colors.append(str(rgb))
                print(f"FAIL: Component 5 — Title text not white. Found colors: {colors}")
        else:
            print(f"FAIL: Component 5 — No paragraphs in document")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
