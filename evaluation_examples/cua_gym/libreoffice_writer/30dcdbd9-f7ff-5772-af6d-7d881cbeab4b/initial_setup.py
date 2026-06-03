"""
Initial Setup: Insert a decorative shape behind the title text.
Task ID: writer_rd_036
Domain: libreoffice_writer

Creates a Writer document with a strategic plan title and content.
The title has NO background decoration and NO white text color.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_036'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title (Heading 1, 24pt, centered) ---
    title_para = doc.add_heading('Strategic Plan 2025-2030', level=1)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_para.runs:
        run.font.size = Pt(24)
        # Keep default color (dark/black) -- NO white text
        run.font.color.rgb = None

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    sub_run = subtitle.add_run('Meridian Global Consulting Group')
    sub_run.font.size = Pt(14)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'This strategic plan outlines the key objectives, initiatives, and performance '
        'targets for Meridian Global Consulting Group over the next five years. Our focus '
        'areas include digital transformation advisory, sustainability consulting, and '
        'expansion into emerging markets across Southeast Asia and Latin America.'
    )
    doc.add_paragraph(
        'The plan was developed through extensive stakeholder engagement sessions held '
        'between September and December 2024, involving over 150 senior leaders, client '
        'representatives, and industry experts. It reflects our commitment to driving '
        'measurable impact for our clients while maintaining double-digit revenue growth.'
    )

    # --- Vision & Mission ---
    doc.add_heading('Vision & Mission', level=2)
    doc.add_paragraph(
        'Vision: To be the most trusted partner for organizations navigating complex '
        'transformations in an increasingly digital and sustainable world.'
    )
    doc.add_paragraph(
        'Mission: We empower enterprises to achieve operational excellence and strategic '
        'agility through data-driven insights, cross-functional expertise, and tailored '
        'solutions that deliver lasting value.'
    )

    # --- Strategic Priorities ---
    doc.add_heading('Strategic Priorities 2025-2030', level=2)
    priorities = [
        'Digital Transformation & AI Integration — Invest $12.5M annually in AI-powered '
        'consulting tools and proprietary analytics platforms to accelerate client outcomes.',
        'Sustainability & ESG Advisory — Build a dedicated practice of 80+ consultants '
        'specializing in carbon accounting, supply chain sustainability, and ESG reporting frameworks.',
        'Geographic Expansion — Establish offices in Singapore, São Paulo, and Nairobi by '
        '2027, targeting $45M in new market revenue by 2030.',
        'Talent Development & Retention — Launch the Meridian Leadership Academy with a '
        'target of 95% retention for senior consultants through competitive compensation, '
        'equity participation, and continuous learning programs.',
        'Client Experience Innovation — Deploy a unified client portal by Q3 2026, reducing '
        'project onboarding time by 40% and improving Net Promoter Score from 72 to 85+.',
    ]
    for p in priorities:
        doc.add_paragraph(p, style='List Bullet')

    # --- Financial Projections ---
    doc.add_heading('Financial Projections', level=2)
    doc.add_paragraph(
        'Revenue is projected to grow from $128M in 2024 to $215M by 2030, representing '
        'a compound annual growth rate (CAGR) of 9.1%. Operating margins are expected to '
        'improve from 18.3% to 22.7% through process automation and economies of scale.'
    )

    # Financial projections table
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    headers = ['Year', 'Revenue ($M)', 'EBITDA ($M)', 'Headcount', 'New Clients']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['2025', '142.0', '28.4', '1,250', '38'],
        ['2026', '156.8', '33.0', '1,420', '45'],
        ['2027', '170.5', '38.2', '1,580', '52'],
        ['2028', '185.0', '42.6', '1,720', '48'],
        ['2029', '200.3', '47.1', '1,880', '55'],
        ['2030', '215.0', '48.8', '2,050', '60'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Implementation Timeline ---
    doc.add_heading('Implementation Timeline', level=2)
    doc.add_paragraph(
        'Phase 1 (2025): Foundation building — complete technology infrastructure upgrades, '
        'hire 120 new consultants, and launch pilot programs in three priority markets.'
    )
    doc.add_paragraph(
        'Phase 2 (2026-2027): Scaling — expand service offerings, open international offices, '
        'and achieve ISO 27001 and B Corp certifications.'
    )
    doc.add_paragraph(
        'Phase 3 (2028-2030): Maturation — optimize operations for sustained profitability, '
        'deepen client relationships through AI-enhanced delivery, and explore strategic acquisitions.'
    )

    # --- Risk Management ---
    doc.add_heading('Risk Management', level=2)
    doc.add_paragraph(
        'Key risks identified include macroeconomic uncertainty, talent market competition, '
        'regulatory changes in target markets, and technology adoption barriers. A dedicated '
        'risk committee chaired by the CFO will conduct quarterly reviews and maintain a '
        'dynamic risk register with defined mitigation strategies for each identified risk.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
