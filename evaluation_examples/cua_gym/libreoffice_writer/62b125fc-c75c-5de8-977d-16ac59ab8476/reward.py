"""
Reward Script: HR Standard Operating Procedures Manual
Task ID: writer_hr_077
Domain: libreoffice_writer
Scoring:
  C1 (0.15) - Document Control section with revision history and approvers tables
  C2 (0.15) - Master Index table listing all 12 SOPs
  C3 (0.25) - 12 SOPs each with 7 required subsections (Purpose, Scope, Definitions, Procedure, Responsible Parties, Related Documents, Revision Notes)
  C4 (0.15) - Outline-numbered procedure steps (NNN.X format)
  C5 (0.15) - Definition and Responsible Parties tables within SOPs (at least 20 of expected 24)
  C6 (0.15) - Multiple sections with SOP-specific headers showing SOP numbers
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_077'

# The 12 expected SOP identifiers
SOP_IDS = [f"SOP-{str(i).zfill(3)}" for i in range(1, 13)]

# Required subsection headings within each SOP
REQUIRED_SUBSECTIONS = [
    "Purpose",
    "Scope",
    "Definitions",
    "Procedure",
    "Responsible Parties",
    "Related Documents",
    "Revision Notes",
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # =========================================================================
    # Component 1: Document Control section with revision history & approvers
    #              tables (0.15 points)
    # In golden: Table 0 = revision history (cols: Version, Date, Author, Description)
    #            Table 1 = document approvers (cols: Name, Title, Approval Date)
    #            A "Document Control" Heading 1 exists.
    # In initial: 0 tables, no such heading.
    # =========================================================================
    try:
        has_doc_control_heading = False
        for p in doc.paragraphs:
            if p.style and p.style.name == 'Heading 1' and 'document control' in p.text.lower():
                has_doc_control_heading = True
                break

        has_revision_table = False
        has_approvers_table = False
        for table in doc.tables:
            headers = [c.text.strip().lower() for c in table.rows[0].cells]
            if 'version' in headers and 'date' in headers:
                has_revision_table = True
            if 'name' in headers and 'title' in headers and ('approval' in ' '.join(headers)):
                has_approvers_table = True

        if has_doc_control_heading and has_revision_table and has_approvers_table:
            print(f"PASS: Component 1 - Document Control section with both tables (0.15 pts)")
            total_score += 0.15
        elif has_doc_control_heading and (has_revision_table or has_approvers_table):
            print(f"PARTIAL: Component 1 - Document Control heading + 1 table (0.075 pts)")
            total_score += 0.075
        else:
            print(f"FAIL: Component 1 - Missing Document Control section/tables "
                  f"(heading={has_doc_control_heading}, rev_table={has_revision_table}, "
                  f"approvers_table={has_approvers_table})")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # =========================================================================
    # Component 2: Master Index table listing all 12 SOPs (0.15 points)
    # In golden: A heading "Master Index" and a table with 13 rows (header + 12 SOPs)
    #            containing SOP-001 through SOP-012.
    # In initial: No such table exists.
    # =========================================================================
    try:
        has_master_index_heading = False
        for p in doc.paragraphs:
            if p.style and p.style.name == 'Heading 1' and 'master index' in p.text.lower():
                has_master_index_heading = True
                break

        # Find a table that contains SOP numbers
        sops_found_in_index = set()
        for table in doc.tables:
            all_text = ' '.join(c.text for row in table.rows for c in row.cells)
            for sop_id in SOP_IDS:
                if sop_id in all_text:
                    sops_found_in_index.add(sop_id)

        sop_count_in_index = len(sops_found_in_index)

        if has_master_index_heading and sop_count_in_index >= 12:
            print(f"PASS: Component 2 - Master Index with all 12 SOPs (0.15 pts)")
            total_score += 0.15
        elif has_master_index_heading and sop_count_in_index >= 6:
            # Partial: at least half the SOPs in the index
            partial = 0.15 * (sop_count_in_index / 12)
            print(f"PARTIAL: Component 2 - Master Index with {sop_count_in_index}/12 SOPs ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - Missing Master Index "
                  f"(heading={has_master_index_heading}, SOPs in index={sop_count_in_index})")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # =========================================================================
    # Component 3: 12 SOPs each with 7 required subsections (0.25 points)
    # Each SOP should have Heading 2 subsections: Purpose, Scope, Definitions,
    # Procedure, Responsible Parties, Related Documents, Revision Notes.
    # In initial: Only Heading 1 titles exist, no Heading 2 subsections.
    # =========================================================================
    try:
        # Build a map of SOP heading positions
        sop_positions = []
        for i, p in enumerate(doc.paragraphs):
            if p.style and p.style.name == 'Heading 1':
                for sop_id in SOP_IDS:
                    if sop_id in p.text:
                        sop_positions.append((sop_id, i))
                        break

        sops_with_full_subsections = 0
        sops_with_partial_subsections = 0
        total_subsection_hits = 0

        for idx, (sop_id, start_pos) in enumerate(sop_positions):
            # Determine end position (next SOP heading or end of doc)
            if idx + 1 < len(sop_positions):
                end_pos = sop_positions[idx + 1][1]
            else:
                end_pos = len(doc.paragraphs)

            # Collect Heading 2 texts within this SOP range
            h2_texts = []
            for j in range(start_pos + 1, end_pos):
                p = doc.paragraphs[j]
                if p.style and p.style.name == 'Heading 2':
                    h2_texts.append(p.text.lower())

            # Check each required subsection
            found_count = 0
            for subsection in REQUIRED_SUBSECTIONS:
                if any(subsection.lower() in h2 for h2 in h2_texts):
                    found_count += 1

            total_subsection_hits += found_count
            if found_count == 7:
                sops_with_full_subsections += 1
            elif found_count >= 4:
                sops_with_partial_subsections += 1

        # Score: proportional to how many SOPs have all 7 subsections
        # Perfect: 12 SOPs * 7 subsections = 84 subsection hits
        max_hits = 12 * 7
        if total_subsection_hits >= 80:
            # Nearly perfect or perfect
            c3_score = 0.25
        elif total_subsection_hits > 0:
            c3_score = 0.25 * (total_subsection_hits / max_hits)
        else:
            c3_score = 0.0

        if c3_score > 0:
            print(f"PASS: Component 3 - {sops_with_full_subsections} SOPs with full subsections, "
                  f"{total_subsection_hits}/{max_hits} subsection hits ({c3_score:.3f} pts)")
            total_score += c3_score
        else:
            print(f"FAIL: Component 3 - No SOP subsections found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # =========================================================================
    # Component 4: Outline-numbered procedure steps (0.15 points)
    # Golden has steps like "001.1", "002.3", "012.5" etc. in procedure sections.
    # Initial has plain text paragraphs with no outline numbering.
    # =========================================================================
    try:
        # Pattern: NNN.N or NNN.NN at the start of paragraph text
        outline_pattern = re.compile(r'^\d{3}\.\d+\s')
        outline_step_count = 0
        sops_with_outline = set()

        for p in doc.paragraphs:
            text = p.text.strip()
            match = outline_pattern.match(text)
            if match:
                outline_step_count += 1
                # Extract SOP number from the prefix (e.g., "001" -> SOP-001)
                prefix = text[:3]
                sop_num = f"SOP-{prefix}"
                sops_with_outline.add(sop_num)

        num_sops_with_outline = len(sops_with_outline)

        if num_sops_with_outline >= 12 and outline_step_count >= 50:
            print(f"PASS: Component 4 - Outline numbering in {num_sops_with_outline} SOPs, "
                  f"{outline_step_count} total steps (0.15 pts)")
            total_score += 0.15
        elif num_sops_with_outline >= 6:
            partial = 0.15 * (num_sops_with_outline / 12)
            print(f"PARTIAL: Component 4 - Outline numbering in {num_sops_with_outline}/12 SOPs "
                  f"({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - Outline numbering found in only {num_sops_with_outline} SOPs "
                  f"({outline_step_count} total steps)")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # =========================================================================
    # Component 5: Definition and Responsible Parties tables within SOPs (0.15 pts)
    # Golden has 24 intra-SOP tables (2 per SOP: Definitions table + Responsible
    # Parties table), plus 3 document-level tables = 27 total.
    # Initial has 0 tables.
    # We check for tables with headers matching "Term"/"Definition" or
    # "Role"/"Responsibilities" patterns.
    # =========================================================================
    try:
        definition_tables = 0
        responsibility_tables = 0
        for table in doc.tables:
            headers = [c.text.strip().lower() for c in table.rows[0].cells]
            if 'term' in headers and 'definition' in headers:
                definition_tables += 1
            if 'role' in headers and 'responsibilities' in headers:
                responsibility_tables += 1

        intra_sop_tables = definition_tables + responsibility_tables

        if intra_sop_tables >= 20:
            print(f"PASS: Component 5 - {definition_tables} definition + {responsibility_tables} "
                  f"responsibility tables = {intra_sop_tables} (0.15 pts)")
            total_score += 0.15
        elif intra_sop_tables >= 10:
            partial = 0.15 * (intra_sop_tables / 24)
            print(f"PARTIAL: Component 5 - {intra_sop_tables}/24 intra-SOP tables ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 - Only {intra_sop_tables} intra-SOP tables found "
                  f"(definition={definition_tables}, responsibility={responsibility_tables})")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # =========================================================================
    # Component 6: Multiple sections with SOP-specific headers (0.15 points)
    # Golden has 12 sections, each with a header containing the SOP number.
    # Initial has 1 section with no header text.
    # =========================================================================
    try:
        num_sections = len(doc.sections)
        sections_with_sop_header = 0

        for section in doc.sections:
            if section.header and section.header.paragraphs:
                header_text = ' '.join(p.text for p in section.header.paragraphs)
                # Check if any SOP-NNN pattern appears in the header
                if re.search(r'SOP-\d{3}', header_text):
                    sections_with_sop_header += 1

        if num_sections >= 10 and sections_with_sop_header >= 10:
            print(f"PASS: Component 6 - {num_sections} sections, {sections_with_sop_header} "
                  f"with SOP headers (0.15 pts)")
            total_score += 0.15
        elif sections_with_sop_header >= 5:
            partial = 0.15 * (sections_with_sop_header / 12)
            print(f"PARTIAL: Component 6 - {sections_with_sop_header}/12 sections with SOP headers "
                  f"({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 - Only {sections_with_sop_header} sections with SOP headers "
                  f"(total sections: {num_sections})")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.3f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
