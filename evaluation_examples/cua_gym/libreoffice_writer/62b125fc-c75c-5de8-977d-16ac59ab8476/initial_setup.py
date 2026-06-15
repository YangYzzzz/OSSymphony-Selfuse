"""
Initial Setup: HR Standard Operating Procedures manual - plain text format
Task ID: writer_hr_077
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_077'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# 12 HR SOP topics with realistic plain-text content
SOP_DATA = [
    {
        "id": "SOP-001",
        "title": "Employee Onboarding",
        "text": (
            "This procedure covers the onboarding process for new employees joining the organization. "
            "HR coordinators are responsible for initiating the onboarding workflow within two business days of the signed offer letter. "
            "The hiring manager must submit a completed Position Authorization Form before the start date. "
            "New hires receive an orientation packet including the employee handbook, benefits enrollment forms, and IT access request. "
            "Background verification must be completed through Sterling Talent Solutions before the employee's first day. "
            "The onboarding checklist includes workspace setup, badge issuance, payroll enrollment, and mandatory compliance training. "
            "IT provisioning requires a minimum of 5 business days lead time for hardware and software access. "
            "The buddy system pairs each new hire with a tenured employee from the same department for the first 90 days. "
            "Orientation sessions are held on the first and third Monday of each month in Conference Room B. "
            "All onboarding documentation must be uploaded to the HRIS within 10 business days of the start date."
        )
    },
    {
        "id": "SOP-002",
        "title": "Performance Review Process",
        "text": (
            "Annual performance reviews are conducted during Q4 for all full-time employees with at least six months of service. "
            "The review cycle begins on October 1 and all evaluations must be submitted by December 15. "
            "Managers complete the Performance Evaluation Form using the five-point competency scale. "
            "Self-assessments are due two weeks before the scheduled review meeting. "
            "Calibration sessions are held at the department level to ensure consistent rating standards. "
            "The VP of Human Resources facilitates cross-departmental calibration for senior staff. "
            "Performance improvement plans are required for any employee receiving an overall rating below 2.5. "
            "Merit increase recommendations are submitted alongside final evaluations to the Compensation team. "
            "Mid-year check-ins are mandatory and documented using the Interim Review Template. "
            "All performance records are retained in the employee's personnel file for a minimum of seven years."
        )
    },
    {
        "id": "SOP-003",
        "title": "Leave and Absence Management",
        "text": (
            "Employees must submit leave requests through the Workday portal at least five business days in advance for planned absences. "
            "Unplanned absences require notification to the direct supervisor within one hour of the scheduled start time. "
            "Annual leave accrual rates are 15 days for employees with 0-3 years of service, 20 days for 4-7 years, and 25 days for 8+ years. "
            "Sick leave is provided at 10 days per calendar year and does not carry over. "
            "FMLA eligibility requires 12 months of employment and 1,250 hours worked in the preceding 12 months. "
            "Short-term disability coverage begins on the eighth consecutive day of absence. "
            "Managers must approve or deny leave requests within 48 hours of submission. "
            "Bereavement leave provides up to five days for immediate family members and three days for extended family. "
            "Jury duty leave is granted with full pay for the duration of service, with proof of summons required. "
            "All leave balances are reconciled by Payroll on the last business day of each month."
        )
    },
    {
        "id": "SOP-004",
        "title": "Recruitment and Hiring",
        "text": (
            "All open positions must be approved by the department head and the VP of Finance before posting. "
            "Job requisitions are created in the Applicant Tracking System by the recruiting coordinator. "
            "Internal candidates are given priority consideration for the first five business days of posting. "
            "External job postings are distributed to LinkedIn, Indeed, Glassdoor, and industry-specific job boards. "
            "Phone screenings are conducted by the recruiter within three business days of application review. "
            "On-site interviews require a structured panel with at least three interviewers using the standardized scorecard. "
            "Hiring managers must provide written feedback within 24 hours of each interview. "
            "Reference checks include two professional references and one supervisory reference. "
            "Offer letters are generated by the Compensation team and reviewed by Legal before distribution. "
            "Time-to-fill targets are 30 days for standard roles and 60 days for director-level and above."
        )
    },
    {
        "id": "SOP-005",
        "title": "Compensation and Benefits Administration",
        "text": (
            "Salary ranges are reviewed annually by the Compensation team using market data from Mercer and Radford surveys. "
            "Pay grades are structured in 12 bands with midpoint differentials of 15 percent between adjacent bands. "
            "Annual merit increases are budgeted at 3-5 percent of total payroll, distributed based on performance ratings. "
            "Benefits open enrollment runs from November 1 through November 30 each year. "
            "The company contributes 80 percent of medical premium costs for employee-only coverage. "
            "The 401(k) plan includes a 4 percent company match, vesting over a four-year graded schedule. "
            "Equity grants for eligible employees are reviewed by the Compensation Committee quarterly. "
            "Salary adjustments outside the annual cycle require VP-level approval and written justification. "
            "New hire benefits eligibility begins on the first day of the month following the start date. "
            "Payroll is processed bi-weekly on Fridays with direct deposit as the default payment method."
        )
    },
    {
        "id": "SOP-006",
        "title": "Employee Disciplinary Actions",
        "text": (
            "Progressive discipline follows a four-step process: verbal warning, written warning, final warning, and termination. "
            "Verbal warnings are documented in the manager's notes and shared with HR for record-keeping. "
            "Written warnings require the employee's acknowledgment signature and are placed in the personnel file. "
            "Suspensions pending investigation may be imposed for serious misconduct allegations. "
            "The Employee Relations team must be consulted before any written warning or more severe action. "
            "Investigations into alleged misconduct are completed within 10 business days unless circumstances require extension. "
            "Employees have the right to appeal disciplinary actions through the Grievance Resolution Procedure. "
            "Termination decisions require approval from the CHRO and review by Employment Legal Counsel. "
            "All disciplinary documentation is maintained in a secure, restricted-access section of the HRIS. "
            "Managers receive annual training on progressive discipline and documentation best practices."
        )
    },
    {
        "id": "SOP-007",
        "title": "Training and Professional Development",
        "text": (
            "Each department allocates a minimum of 2 percent of payroll budget to employee training and development. "
            "Mandatory compliance training includes anti-harassment, data privacy, and workplace safety modules. "
            "Compliance training must be completed within 30 days of hire and annually thereafter. "
            "Tuition reimbursement covers up to $5,250 per calendar year for pre-approved degree programs. "
            "Professional certification reimbursement is available for industry-recognized credentials approved by the manager. "
            "Individual Development Plans are created during the Q1 performance planning cycle. "
            "The Learning Management System tracks all training completions and certifications. "
            "External conference attendance requires pre-approval and a post-event knowledge-sharing presentation. "
            "Mentorship programs are coordinated by the Talent Development team with quarterly matching cycles. "
            "Leadership development programs are offered to high-potential employees identified through the talent review process."
        )
    },
    {
        "id": "SOP-008",
        "title": "Workplace Health and Safety",
        "text": (
            "The Safety Committee meets monthly to review incident reports and inspect workplace conditions. "
            "All workplace injuries must be reported to the Safety Officer within 24 hours of occurrence. "
            "First aid kits are inspected monthly and restocked by the Facilities team. "
            "Fire drills are conducted quarterly with evacuation routes posted on each floor. "
            "Ergonomic assessments are available upon request through the Employee Wellness program. "
            "Personal protective equipment is provided at no cost for roles requiring safety gear. "
            "OSHA 300 logs are maintained and posted annually from February 1 through April 30. "
            "Hazardous material handling requires completion of the HAZCOM training module. "
            "Return-to-work programs coordinate with occupational health providers for modified duty assignments. "
            "Annual safety training covers emergency procedures, fire safety, and active threat response."
        )
    },
    {
        "id": "SOP-009",
        "title": "Employee Offboarding and Separation",
        "text": (
            "Voluntary resignations require a minimum two-week written notice submitted to the direct supervisor and HR. "
            "Exit interviews are conducted by HR within the employee's final five business days. "
            "IT access revocation occurs at 5:00 PM on the employee's last day for voluntary separations. "
            "For involuntary terminations, IT access is revoked immediately upon notification. "
            "The final paycheck includes all accrued but unused vacation days per state law requirements. "
            "Company property including laptops, badges, parking passes, and keys must be returned on the last day. "
            "COBRA notification is sent within 14 days of the qualifying event by the Benefits administrator. "
            "Knowledge transfer plans are developed by the departing employee and approved by the manager. "
            "Separation agreements for negotiated departures are drafted by Employment Legal Counsel. "
            "Personnel files for separated employees are archived and retained for seven years."
        )
    },
    {
        "id": "SOP-010",
        "title": "Equal Employment Opportunity and Anti-Discrimination",
        "text": (
            "The company maintains a zero-tolerance policy for discrimination based on race, gender, age, disability, religion, national origin, or other protected characteristics. "
            "All employment decisions including hiring, promotion, and compensation are based solely on merit and qualifications. "
            "Discrimination complaints are filed with the Employee Relations team via the confidential reporting hotline or online portal. "
            "Investigations are initiated within 48 hours of receipt of a formal complaint. "
            "Reasonable accommodations under the ADA are coordinated through the HR Business Partner and Facilities. "
            "Affirmative Action Plans are updated annually and reviewed by the Director of Diversity and Inclusion. "
            "EEO-1 reports are filed annually with the EEOC by the September 30 deadline. "
            "Anti-harassment training is mandatory for all employees annually and for managers upon promotion. "
            "Retaliation against employees who report discrimination in good faith is strictly prohibited. "
            "The Chief Diversity Officer provides quarterly metrics on workforce demographics and complaint trends."
        )
    },
    {
        "id": "SOP-011",
        "title": "Remote Work and Flexible Arrangements",
        "text": (
            "Eligible employees may request remote work arrangements through the Flexible Work Agreement Form. "
            "Hybrid schedules require a minimum of three in-office days per week unless otherwise approved. "
            "Remote workers must maintain a dedicated workspace meeting ergonomic and security standards. "
            "VPN connectivity and multi-factor authentication are required for all remote access to company systems. "
            "Managers conduct monthly one-on-one meetings with remote team members to assess engagement and performance. "
            "Home office equipment stipends of up to $500 are provided for approved full-time remote employees. "
            "Remote work agreements are reviewed semi-annually and may be modified based on business needs. "
            "Time zone considerations require core overlap hours of 10:00 AM to 3:00 PM Eastern for distributed teams. "
            "Data handling in remote environments must comply with the Information Security Policy. "
            "International remote work requires prior approval from Legal, Tax, and HR due to cross-border compliance requirements."
        )
    },
    {
        "id": "SOP-012",
        "title": "Employee Records Management",
        "text": (
            "Employee records are maintained in the HRIS with access restricted to authorized HR personnel. "
            "Personnel files include employment applications, offer letters, performance evaluations, and disciplinary records. "
            "Medical records are stored separately from personnel files in compliance with HIPAA and ADA requirements. "
            "I-9 forms are retained for three years after the date of hire or one year after termination, whichever is later. "
            "Records retention schedules are reviewed annually by the Records Management Coordinator and Legal. "
            "Employees may request access to their personnel file with 48 hours advance notice. "
            "File audits are conducted quarterly to verify completeness and compliance with retention policies. "
            "Digital records are backed up daily to encrypted cloud storage with 256-bit AES encryption. "
            "Physical records are stored in locked cabinets within the HR suite with key-card access. "
            "Disposal of expired records follows the secure destruction protocol using certified shredding services."
        )
    },
]


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title - plain text, minimal formatting
    doc.add_heading('HR Standard Operating Procedures', level=0)
    doc.add_paragraph('Acme Global Corporation')
    doc.add_paragraph('Human Resources Department')
    doc.add_paragraph('')

    # Plain text list of the 12 SOPs - no structured formatting
    for sop in SOP_DATA:
        doc.add_heading(f'{sop["id"]} - {sop["title"]}', level=1)
        doc.add_paragraph(sop["text"])
        doc.add_paragraph('')  # spacer

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
