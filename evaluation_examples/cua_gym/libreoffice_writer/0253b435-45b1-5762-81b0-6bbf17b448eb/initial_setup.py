"""
Initial Setup: Multi-level numbered list from plain text outline
Task ID: wrpara_035
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'wrpara_035'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Project Work Breakdown Structure', level=1)

    # Introductory paragraph
    doc.add_paragraph(
        'The following outline details the work breakdown structure for the '
        'Enterprise Resource Planning (ERP) system migration project at Meridian Technologies. '
        'All items are currently listed as plain text and need to be converted into a '
        'properly formatted multi-level numbered list.'
    )

    doc.add_paragraph('')  # blank line separator

    # WBS items as plain text with manual space indentation
    # Level 1 items have no indentation
    # Level 2 items have 4 spaces
    # Level 3 items have 8 spaces

    wbs_items = [
        # Level 1 - Item 1
        ("Project Planning and Initiation", 0),
        ("    Stakeholder Requirements Gathering", 1),
        ("        Conduct Executive Interviews", 2),
        ("        Distribute Department Surveys", 2),
        ("    Feasibility Study and Risk Assessment", 1),
        ("    Project Charter and Timeline Development", 1),

        # Level 1 - Item 2
        ("System Design and Architecture", 0),
        ("    Database Schema Design", 1),
        ("        Define Entity Relationships", 2),
        ("        Establish Data Migration Mappings", 2),
        ("    User Interface Wireframing", 1),
        ("    API Integration Planning", 1),

        # Level 1 - Item 3
        ("Development and Implementation", 0),
        ("    Core Module Development", 1),
        ("        Build Inventory Management Module", 2),
        ("        Build Financial Reporting Module", 2),
        ("    Third-Party Integration Development", 1),
        ("        Connect Payment Gateway Services", 2),
        ("        Integrate Cloud Storage Providers", 2),
        ("    Quality Assurance and Testing", 1),

        # Level 1 - Item 4
        ("Deployment and Post-Launch Support", 0),
        ("    Staging Environment Deployment", 1),
        ("    Production Rollout and Monitoring", 1),
        ("        Configure Performance Dashboards", 2),
        ("        Establish Incident Response Procedures", 2),
        ("    End-User Training and Documentation", 1),
    ]

    for text, level in wbs_items:
        # Add as plain text paragraphs (no list style)
        para = doc.add_paragraph(text)
        # No list formatting - just plain text with leading spaces for indentation

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
