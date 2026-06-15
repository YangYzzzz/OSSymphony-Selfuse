"""
Initial Setup: Apply bold formatting to all Heading 1 and Heading 2 paragraphs
Task ID: osworld_writer_easy_006
Domain: libreoffice_writer

Creates an annual report draft with 3 Heading 1 sections and 6 Heading 2 subsections.
Some headings are bold, some are NOT bold (mixed state) — agent must apply bold to all.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_easy_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_run(para, text, bold=None):
    """Add a run to a heading paragraph with optional explicit bold setting."""
    run = para.add_run(text)
    if bold is not None:
        run.bold = bold
    return run


def create_initial():
    doc = Document()

    # --- Title (not a heading style, just a styled normal para) ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("Meridian Technologies Group")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run("Annual Report 2024")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    doc.add_paragraph()  # blank line

    # =====================================================================
    # HEADING 1 — Section 1: bold=True (already bold)
    # =====================================================================
    h1_1 = doc.add_heading('', level=1)
    h1_1.clear()
    add_heading_run(h1_1, 'Executive Summary', bold=True)

    doc.add_paragraph(
        "Meridian Technologies Group concluded fiscal year 2024 with record-breaking revenue "
        "of $2.4 billion, representing a 17% increase over the prior year. Our strategic "
        "investments in cloud infrastructure and AI-driven product lines have positioned us "
        "for continued growth in the coming decade."
    )
    doc.add_paragraph(
        "This report details operational performance, financial results, and forward-looking "
        "strategic priorities across our three core business segments: Enterprise Solutions, "
        "Consumer Products, and Professional Services."
    )

    # HEADING 2 — 1.1: bold=False (NOT bold — agent must fix)
    h2_1 = doc.add_heading('', level=2)
    h2_1.clear()
    add_heading_run(h2_1, 'Financial Highlights', bold=False)

    doc.add_paragraph(
        "Total revenue reached $2.4B, with gross margins improving to 68.3% from 65.1% "
        "in 2023. Operating income was $612M, and diluted earnings per share grew 22% "
        "year-over-year to $4.87. The company returned $340M to shareholders through "
        "dividends and share buybacks."
    )

    # HEADING 2 — 1.2: bold=True (already bold)
    h2_2 = doc.add_heading('', level=2)
    h2_2.clear()
    add_heading_run(h2_2, 'Key Milestones', bold=True)

    doc.add_paragraph(
        "During 2024, Meridian completed the acquisition of DataSync Corp for $180M, "
        "expanded into four new international markets, launched the Meridian Cloud Platform "
        "version 3.0, and achieved ISO 27001 certification across all business units."
    )

    # =====================================================================
    # HEADING 1 — Section 2: bold=False (NOT bold — agent must fix)
    # =====================================================================
    h1_2 = doc.add_heading('', level=1)
    h1_2.clear()
    add_heading_run(h1_2, 'Business Segment Performance', bold=False)

    doc.add_paragraph(
        "Each of our three core business segments delivered strong performance in 2024. "
        "Enterprise Solutions continued to be our largest revenue driver, while Consumer "
        "Products posted the highest growth rate at 31%. Professional Services maintained "
        "high margin performance with stable client retention."
    )

    # HEADING 2 — 2.1: bold=False (NOT bold — agent must fix)
    h2_3 = doc.add_heading('', level=2)
    h2_3.clear()
    add_heading_run(h2_3, 'Enterprise Solutions', bold=False)

    doc.add_paragraph(
        "Enterprise Solutions revenue totaled $1.1B, a 12% increase year-over-year. "
        "Key contracts with Fortune 500 clients including Apex Financial Group, Northridge "
        "Healthcare Network, and Pacific Logistics Corp drove growth. Annual recurring "
        "revenue (ARR) from enterprise cloud subscriptions reached $650M."
    )

    # HEADING 2 — 2.2: bold=True (already bold)
    h2_4 = doc.add_heading('', level=2)
    h2_4.clear()
    add_heading_run(h2_4, 'Consumer Products', bold=True)

    doc.add_paragraph(
        "Consumer Products segment surged 31% to $820M in revenue. The Meridian Home Hub "
        "smart device ecosystem grew to 4.2 million active users. The launch of MeridianOS "
        "for personal devices attracted strong developer adoption with 18,000 registered "
        "third-party app developers by year end."
    )

    # HEADING 2 — 2.3: bold=False (NOT bold — agent must fix)
    h2_5 = doc.add_heading('', level=2)
    h2_5.clear()
    add_heading_run(h2_5, 'Professional Services', bold=False)

    doc.add_paragraph(
        "Professional Services contributed $480M in revenue with a 41% operating margin. "
        "Client satisfaction scores reached an all-time high of 4.6/5.0. The division "
        "added 127 new enterprise consulting engagements and maintained a 94% client "
        "retention rate across existing contracts."
    )

    # =====================================================================
    # HEADING 1 — Section 3: bold=True (already bold)
    # =====================================================================
    h1_3 = doc.add_heading('', level=1)
    h1_3.clear()
    add_heading_run(h1_3, 'Strategic Outlook and Future Priorities', bold=True)

    doc.add_paragraph(
        "Looking ahead to 2025, Meridian Technologies Group is committed to accelerating "
        "growth through targeted investments in artificial intelligence, expanded global "
        "partnerships, and next-generation product development. We anticipate total revenue "
        "of $2.8–3.0B for the upcoming fiscal year."
    )

    # HEADING 2 — 3.1: bold=False (NOT bold — agent must fix)
    h2_6 = doc.add_heading('', level=2)
    h2_6.clear()
    add_heading_run(h2_6, 'Innovation and R&D Investment', bold=False)

    doc.add_paragraph(
        "We will increase R&D spending by 25% to $380M in 2025, with priority focus on "
        "generative AI integration across our product suite, quantum-ready encryption for "
        "enterprise clients, and advanced analytics capabilities within the Meridian Cloud "
        "Platform. Our innovation centers in Singapore, Toronto, and Dublin will expand "
        "headcount by 400 engineers."
    )

    # HEADING 2 — 3.2: bold=True (already bold)
    h2_7 = doc.add_heading('', level=2)
    h2_7.clear()
    add_heading_run(h2_7, 'Global Expansion Strategy', bold=True)

    doc.add_paragraph(
        "Meridian will enter five additional markets in 2025: Brazil, South Africa, "
        "Vietnam, Poland, and the United Arab Emirates. Localized product offerings and "
        "regional partnership agreements are already in progress. We project international "
        "revenue to grow from 34% to 45% of total revenue by end of 2026."
    )

    doc.add_paragraph(
        "Meridian Technologies Group remains committed to delivering long-term value for "
        "shareholders, clients, and employees alike. We thank our global team of 12,400 "
        "employees for their dedication and our investors for their continued confidence."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
