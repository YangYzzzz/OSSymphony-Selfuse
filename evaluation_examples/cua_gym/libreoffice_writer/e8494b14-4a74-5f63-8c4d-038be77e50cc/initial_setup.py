"""
Initial Setup: Create Contract_v1.docx and Contract_v2.docx for document comparison task
Task ID: writer_pd_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_040'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

def add_page_break(doc):
    """Add a page break to the document."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)

def create_contract_v1():
    """Create the base contract document (Contract_v1.docx) - approximately 3 pages."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ===== PAGE 1 =====
    # Title
    title = doc.add_heading('SERVICE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Preamble
    p = doc.add_paragraph()
    p.add_run('Agreement Number: ').bold = True
    p.add_run('SA-2025-0847')

    p = doc.add_paragraph()
    p.add_run('Effective Date: ').bold = True
    p.add_run('March 15, 2025')

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.add_run('This Service Agreement ("Agreement") is entered into by and between ')
    p.add_run('Meridian Technologies Inc.').bold = True
    p.add_run(' ("Service Provider"), a corporation organized under the laws of the State of Delaware, '
              'with its principal office at 1200 Innovation Drive, Suite 400, Austin, TX 78701, and ')
    p.add_run('Cascade Financial Group LLC').bold = True
    p.add_run(' ("Client"), a limited liability company organized under the laws of the State of New York, '
              'with its principal office at 350 Park Avenue, 22nd Floor, New York, NY 10022.')

    doc.add_paragraph()

    # Section 1
    doc.add_heading('1. SCOPE OF SERVICES', level=1)

    doc.add_paragraph(
        '1.1 The Service Provider shall deliver comprehensive IT infrastructure management services '
        'as outlined in Exhibit A attached hereto. These services include, but are not limited to, '
        'network monitoring, server maintenance, cybersecurity assessments, and help desk support.'
    )

    doc.add_paragraph(
        '1.2 The Service Provider shall assign a dedicated team of no fewer than five (5) qualified '
        'professionals to perform the services described in this Agreement. The team shall include '
        'a project manager, two senior engineers, and two support specialists.'
    )

    doc.add_paragraph(
        '1.3 All services shall be performed in accordance with industry best practices and applicable '
        'regulatory requirements, including but not limited to SOC 2 Type II compliance standards.'
    )

    # Section 2
    doc.add_heading('2. COMPENSATION AND PAYMENT TERMS', level=1)

    doc.add_paragraph(
        '2.1 The Client shall pay the Service Provider a monthly retainer fee of Forty-Five Thousand '
        'Dollars ($45,000) for the services described in Section 1, payable within fifteen (15) business '
        'days of receipt of invoice.'
    )

    doc.add_paragraph(
        '2.2 Additional services requested beyond the scope defined in Exhibit A shall be billed at '
        'the following hourly rates: Senior Engineer - $185/hour; Project Manager - $200/hour; '
        'Support Specialist - $125/hour.'
    )

    doc.add_paragraph(
        '2.3 Late payments shall accrue interest at a rate of one and a half percent (1.5%) per month '
        'on the outstanding balance, compounded monthly.'
    )

    # ===== PAGE 2 =====
    add_page_break(doc)

    doc.add_heading('3. TERM AND RENEWAL', level=1)

    doc.add_paragraph(
        '3.1 This Agreement shall commence on the Effective Date and continue for an initial term of '
        'twenty-four (24) months ("Initial Term").'
    )

    doc.add_paragraph(
        '3.2 Upon expiration of the Initial Term, this Agreement shall automatically renew for '
        'successive twelve (12) month periods ("Renewal Terms") unless either party provides written '
        'notice of non-renewal at least sixty (60) days prior to the expiration of the then-current term.'
    )

    doc.add_paragraph(
        '3.3 During any Renewal Term, the monthly retainer fee may be adjusted by no more than five '
        'percent (5%) upon sixty (60) days written notice by the Service Provider.'
    )

    doc.add_heading('4. CONFIDENTIALITY', level=1)

    doc.add_paragraph(
        '4.1 Each party acknowledges that in the course of performing its obligations under this '
        'Agreement, it may receive or have access to confidential and proprietary information of the '
        'other party ("Confidential Information").'
    )

    doc.add_paragraph(
        '4.2 Confidential Information shall include, without limitation, trade secrets, business plans, '
        'financial data, customer lists, technical specifications, and any information marked as '
        '"Confidential" or that a reasonable person would understand to be confidential.'
    )

    doc.add_paragraph(
        '4.3 The receiving party shall protect Confidential Information using the same degree of care '
        'it uses to protect its own confidential information, but in no event less than reasonable care. '
        'The obligations under this section shall survive termination for a period of three (3) years.'
    )

    doc.add_heading('5. INTELLECTUAL PROPERTY', level=1)

    doc.add_paragraph(
        '5.1 All pre-existing intellectual property shall remain the sole property of the party that '
        'owned it prior to this Agreement.'
    )

    doc.add_paragraph(
        '5.2 Any work product, deliverables, or materials created by the Service Provider specifically '
        'for the Client under this Agreement ("Work Product") shall be owned by the Client upon full '
        'payment of all applicable fees.'
    )

    # ===== PAGE 3 =====
    add_page_break(doc)

    doc.add_heading('6. TERMINATION', level=1)

    doc.add_paragraph(
        '6.1 Either party may terminate this Agreement for cause upon written notice if the other party '
        'materially breaches any provision of this Agreement and fails to cure such breach within '
        'thirty (30) days of receiving written notice specifying the breach.'
    )

    doc.add_paragraph(
        '6.2 Either party may terminate this Agreement without cause by providing the other party '
        'with thirty (30) days prior written notice.'
    )

    doc.add_paragraph(
        '6.3 Upon termination, the Service Provider shall deliver to the Client all Work Product, '
        'Confidential Information, and any other materials belonging to the Client within fifteen (15) '
        'business days of the effective date of termination.'
    )

    doc.add_heading('7. LIABILITY AND INDEMNIFICATION', level=1)

    doc.add_paragraph(
        '7.1 The Service Provider\'s total aggregate liability under this Agreement shall not exceed '
        'the total fees paid by the Client during the twelve (12) month period immediately preceding '
        'the event giving rise to the claim.'
    )

    doc.add_paragraph(
        '7.2 Neither party shall be liable for any indirect, incidental, special, consequential, or '
        'punitive damages arising out of or related to this Agreement, regardless of the theory of '
        'liability.'
    )

    doc.add_paragraph(
        '7.3 Each party shall indemnify, defend, and hold harmless the other party from and against '
        'any third-party claims arising from the indemnifying party\'s negligence or willful misconduct '
        'in connection with this Agreement.'
    )

    doc.add_heading('8. GENERAL PROVISIONS', level=1)

    doc.add_paragraph(
        '8.1 This Agreement constitutes the entire agreement between the parties with respect to the '
        'subject matter hereof and supersedes all prior agreements, negotiations, and discussions.'
    )

    doc.add_paragraph(
        '8.2 This Agreement shall be governed by and construed in accordance with the laws of the '
        'State of New York, without regard to its conflict of laws provisions.'
    )

    doc.add_paragraph(
        '8.3 Any disputes arising under or in connection with this Agreement shall be resolved through '
        'binding arbitration administered by the American Arbitration Association in New York, New York.'
    )

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('IN WITNESS WHEREOF').bold = True
    p.add_run(', the parties have executed this Agreement as of the date first written above.')

    doc.add_paragraph()

    # Signature lines
    for party, name, title_text in [
        ('Meridian Technologies Inc.', 'Jonathan R. Whitfield', 'Chief Executive Officer'),
        ('Cascade Financial Group LLC', 'Alexandra M. Thornton', 'Managing Director'),
    ]:
        p = doc.add_paragraph()
        p.add_run('_' * 40)
        p = doc.add_paragraph()
        p.add_run(f'{name}').bold = True
        p = doc.add_paragraph(f'{title_text}')
        p = doc.add_paragraph(f'{party}')
        doc.add_paragraph()

    v1_path = f'{WORKDIR}/Contract_v1.docx'
    doc.save(v1_path)
    print(f'Contract_v1.docx created: {v1_path}')
    return doc

def create_contract_v2():
    """
    Create Contract_v2.docx with 12 differences from v1:
    - 8 formatting changes
    - 4 content changes (including the notice period change on page 3)
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ===== PAGE 1 =====
    # Title - FORMATTING CHANGE 1: Title font size changed
    title = doc.add_heading('SERVICE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Preamble
    p = doc.add_paragraph()
    p.add_run('Agreement Number: ').bold = True
    p.add_run('SA-2025-0847')

    p = doc.add_paragraph()
    p.add_run('Effective Date: ').bold = True
    p.add_run('March 15, 2025')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('This Service Agreement ("Agreement") is entered into by and between ')
    # FORMATTING CHANGE 2: Company name in italic as well as bold
    run = p.add_run('Meridian Technologies Inc.')
    run.bold = True
    run.italic = True
    p.add_run(' ("Service Provider"), a corporation organized under the laws of the State of Delaware, '
              'with its principal office at 1200 Innovation Drive, Suite 400, Austin, TX 78701, and ')
    # FORMATTING CHANGE 3: Second company name in italic as well as bold
    run = p.add_run('Cascade Financial Group LLC')
    run.bold = True
    run.italic = True
    p.add_run(' ("Client"), a limited liability company organized under the laws of the State of New York, '
              'with its principal office at 350 Park Avenue, 22nd Floor, New York, NY 10022.')

    doc.add_paragraph()

    # Section 1
    doc.add_heading('1. SCOPE OF SERVICES', level=1)

    doc.add_paragraph(
        '1.1 The Service Provider shall deliver comprehensive IT infrastructure management services '
        'as outlined in Exhibit A attached hereto. These services include, but are not limited to, '
        'network monitoring, server maintenance, cybersecurity assessments, and help desk support.'
    )

    # CONTENT CHANGE 1: Changed from "five (5)" to "six (6)"
    doc.add_paragraph(
        '1.2 The Service Provider shall assign a dedicated team of no fewer than six (6) qualified '
        'professionals to perform the services described in this Agreement. The team shall include '
        'a project manager, two senior engineers, and three support specialists.'
    )

    # FORMATTING CHANGE 4: Added bold to "SOC 2 Type II"
    p = doc.add_paragraph(
        '1.3 All services shall be performed in accordance with industry best practices and applicable '
        'regulatory requirements, including but not limited to '
    )
    run = p.add_run('SOC 2 Type II')
    run.bold = True
    p.add_run(' compliance standards.')

    # Section 2
    doc.add_heading('2. COMPENSATION AND PAYMENT TERMS', level=1)

    # FORMATTING CHANGE 5: Dollar amount underlined
    p = doc.add_paragraph(
        '2.1 The Client shall pay the Service Provider a monthly retainer fee of '
    )
    run = p.add_run('Forty-Five Thousand Dollars ($45,000)')
    run.underline = True
    p.add_run(
        ' for the services described in Section 1, payable within fifteen (15) business '
        'days of receipt of invoice.'
    )

    # CONTENT CHANGE 2: Changed hourly rates
    doc.add_paragraph(
        '2.2 Additional services requested beyond the scope defined in Exhibit A shall be billed at '
        'the following hourly rates: Senior Engineer - $195/hour; Project Manager - $210/hour; '
        'Support Specialist - $135/hour.'
    )

    doc.add_paragraph(
        '2.3 Late payments shall accrue interest at a rate of one and a half percent (1.5%) per month '
        'on the outstanding balance, compounded monthly.'
    )

    # ===== PAGE 2 =====
    add_page_break(doc)

    doc.add_heading('3. TERM AND RENEWAL', level=1)

    doc.add_paragraph(
        '3.1 This Agreement shall commence on the Effective Date and continue for an initial term of '
        'twenty-four (24) months ("Initial Term").'
    )

    # FORMATTING CHANGE 6: "sixty (60) days" in bold
    p = doc.add_paragraph(
        '3.2 Upon expiration of the Initial Term, this Agreement shall automatically renew for '
        'successive twelve (12) month periods ("Renewal Terms") unless either party provides written '
        'notice of non-renewal at least '
    )
    run = p.add_run('sixty (60) days')
    run.bold = True
    p.add_run(' prior to the expiration of the then-current term.')

    doc.add_paragraph(
        '3.3 During any Renewal Term, the monthly retainer fee may be adjusted by no more than five '
        'percent (5%) upon sixty (60) days written notice by the Service Provider.'
    )

    doc.add_heading('4. CONFIDENTIALITY', level=1)

    doc.add_paragraph(
        '4.1 Each party acknowledges that in the course of performing its obligations under this '
        'Agreement, it may receive or have access to confidential and proprietary information of the '
        'other party ("Confidential Information").'
    )

    doc.add_paragraph(
        '4.2 Confidential Information shall include, without limitation, trade secrets, business plans, '
        'financial data, customer lists, technical specifications, and any information marked as '
        '"Confidential" or that a reasonable person would understand to be confidential.'
    )

    # FORMATTING CHANGE 7: Changed "three (3) years" to italic
    p = doc.add_paragraph(
        '4.3 The receiving party shall protect Confidential Information using the same degree of care '
        'it uses to protect its own confidential information, but in no event less than reasonable care. '
        'The obligations under this section shall survive termination for a period of '
    )
    run = p.add_run('three (3) years')
    run.italic = True
    p.add_run('.')

    doc.add_heading('5. INTELLECTUAL PROPERTY', level=1)

    doc.add_paragraph(
        '5.1 All pre-existing intellectual property shall remain the sole property of the party that '
        'owned it prior to this Agreement.'
    )

    # FORMATTING CHANGE 8: "Work Product" in bold and underlined
    p = doc.add_paragraph(
        '5.2 Any work product, deliverables, or materials created by the Service Provider specifically '
        'for the Client under this Agreement ("'
    )
    run = p.add_run('Work Product')
    run.bold = True
    run.underline = True
    p.add_run('") shall be owned by the Client upon full payment of all applicable fees.')

    # ===== PAGE 3 =====
    add_page_break(doc)

    doc.add_heading('6. TERMINATION', level=1)

    doc.add_paragraph(
        '6.1 Either party may terminate this Agreement for cause upon written notice if the other party '
        'materially breaches any provision of this Agreement and fails to cure such breach within '
        'thirty (30) days of receiving written notice specifying the breach.'
    )

    # CONTENT CHANGE 3 (THE KEY ONE): Changed "thirty (30) days" to "fifteen (15) days"
    # This is the change on page 3, paragraph 2 of section 6 that must be REJECTED
    doc.add_paragraph(
        '6.2 Either party may terminate this Agreement without cause by providing the other party '
        'with fifteen (15) days prior written notice.'
    )

    # CONTENT CHANGE 4: Changed "fifteen (15)" to "ten (10)" business days
    doc.add_paragraph(
        '6.3 Upon termination, the Service Provider shall deliver to the Client all Work Product, '
        'Confidential Information, and any other materials belonging to the Client within ten (10) '
        'business days of the effective date of termination.'
    )

    doc.add_heading('7. LIABILITY AND INDEMNIFICATION', level=1)

    doc.add_paragraph(
        '7.1 The Service Provider\'s total aggregate liability under this Agreement shall not exceed '
        'the total fees paid by the Client during the twelve (12) month period immediately preceding '
        'the event giving rise to the claim.'
    )

    doc.add_paragraph(
        '7.2 Neither party shall be liable for any indirect, incidental, special, consequential, or '
        'punitive damages arising out of or related to this Agreement, regardless of the theory of '
        'liability.'
    )

    doc.add_paragraph(
        '7.3 Each party shall indemnify, defend, and hold harmless the other party from and against '
        'any third-party claims arising from the indemnifying party\'s negligence or willful misconduct '
        'in connection with this Agreement.'
    )

    doc.add_heading('8. GENERAL PROVISIONS', level=1)

    doc.add_paragraph(
        '8.1 This Agreement constitutes the entire agreement between the parties with respect to the '
        'subject matter hereof and supersedes all prior agreements, negotiations, and discussions.'
    )

    doc.add_paragraph(
        '8.2 This Agreement shall be governed by and construed in accordance with the laws of the '
        'State of New York, without regard to its conflict of laws provisions.'
    )

    doc.add_paragraph(
        '8.3 Any disputes arising under or in connection with this Agreement shall be resolved through '
        'binding arbitration administered by the American Arbitration Association in New York, New York.'
    )

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('IN WITNESS WHEREOF').bold = True
    p.add_run(', the parties have executed this Agreement as of the date first written above.')

    doc.add_paragraph()

    for party, name, title_text in [
        ('Meridian Technologies Inc.', 'Jonathan R. Whitfield', 'Chief Executive Officer'),
        ('Cascade Financial Group LLC', 'Alexandra M. Thornton', 'Managing Director'),
    ]:
        p = doc.add_paragraph()
        p.add_run('_' * 40)
        p = doc.add_paragraph()
        p.add_run(f'{name}').bold = True
        p = doc.add_paragraph(f'{title_text}')
        p = doc.add_paragraph(f'{party}')
        doc.add_paragraph()

    v2_path = f'{WORKDIR}/Contract_v2.docx'
    doc.save(v2_path)
    print(f'Contract_v2.docx created: {v2_path}')


def create_initial():
    create_contract_v1()
    create_contract_v2()

    # GUI-ready startup: open Contract_v1.docx in Writer
    launch_gui(f'libreoffice --writer "{WORKDIR}/Contract_v1.docx"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with Contract_v1.docx on DISPLAY=:0')


create_initial()
