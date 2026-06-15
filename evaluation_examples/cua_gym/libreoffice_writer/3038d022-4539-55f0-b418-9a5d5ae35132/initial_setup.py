"""
Initial Setup: Mail merge contract template with data source
Task ID: writer_mt_029
Domain: libreoffice_writer
"""

import os
import csv
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_029'
TEMPLATE_FILE = f'{WORKDIR}/Contract_Template.docx'
DATA_FILE = f'{WORKDIR}/Contracts.csv'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# --- 15 contract records ---
CONTRACT_RECORDS = [
    {"PartyName": "Meridian Technologies Inc.", "ContractValue": "$245,000", "EffectiveDate": "January 15, 2025", "TermLength": "24 months"},
    {"PartyName": "Cascade Financial Group", "ContractValue": "$187,500", "EffectiveDate": "February 1, 2025", "TermLength": "12 months"},
    {"PartyName": "Vertex Software Solutions", "ContractValue": "$312,000", "EffectiveDate": "March 10, 2025", "TermLength": "36 months"},
    {"PartyName": "Horizon Healthcare Partners", "ContractValue": "$156,800", "EffectiveDate": "April 5, 2025", "TermLength": "18 months"},
    {"PartyName": "Atlas Construction LLC", "ContractValue": "$425,000", "EffectiveDate": "May 20, 2025", "TermLength": "24 months"},
    {"PartyName": "Pinnacle Marketing Agency", "ContractValue": "$98,750", "EffectiveDate": "June 1, 2025", "TermLength": "12 months"},
    {"PartyName": "Silverline Logistics Corp.", "ContractValue": "$267,300", "EffectiveDate": "July 15, 2025", "TermLength": "36 months"},
    {"PartyName": "Redwood Environmental Services", "ContractValue": "$134,200", "EffectiveDate": "August 8, 2025", "TermLength": "18 months"},
    {"PartyName": "Quantum Research Labs", "ContractValue": "$389,000", "EffectiveDate": "September 12, 2025", "TermLength": "24 months"},
    {"PartyName": "Bluecrest Consulting Group", "ContractValue": "$211,500", "EffectiveDate": "October 1, 2025", "TermLength": "12 months"},
    {"PartyName": "Oakbridge Manufacturing", "ContractValue": "$478,600", "EffectiveDate": "November 3, 2025", "TermLength": "36 months"},
    {"PartyName": "Summit Legal Advisors", "ContractValue": "$145,000", "EffectiveDate": "December 15, 2025", "TermLength": "18 months"},
    {"PartyName": "Trident Aerospace Systems", "ContractValue": "$562,400", "EffectiveDate": "January 8, 2026", "TermLength": "24 months"},
    {"PartyName": "Evergreen Energy Solutions", "ContractValue": "$203,750", "EffectiveDate": "February 14, 2026", "TermLength": "12 months"},
    {"PartyName": "Lighthouse Data Analytics", "ContractValue": "$176,900", "EffectiveDate": "March 22, 2026", "TermLength": "36 months"},
]


def create_template():
    """Create a professional contract template with merge fields."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('SERVICE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Blank line
    doc.add_paragraph('')

    # Date line
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run('Effective Date: <EffectiveDate>')
    run.font.size = Pt(11)

    doc.add_paragraph('')

    # Parties section
    p = doc.add_paragraph()
    run = p.add_run('PARTIES')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('This Service Agreement ("Agreement") is entered into by and between ')
    run.font.size = Pt(11)
    run = p.add_run('Global Enterprises Corp.')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(' ("Provider") and ')
    run.font.size = Pt(11)
    run = p.add_run('<PartyName>')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(' ("Client"), collectively referred to as the "Parties".')
    run.font.size = Pt(11)

    doc.add_paragraph('')

    # Scope section
    p = doc.add_paragraph()
    run = p.add_run('1. SCOPE OF SERVICES')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('The Provider agrees to deliver professional consulting and technology services as outlined in the attached Statement of Work. The scope encompasses strategic planning, implementation support, and ongoing maintenance for the duration specified herein.')
    run.font.size = Pt(11)

    doc.add_paragraph('')

    # Compensation section
    p = doc.add_paragraph()
    run = p.add_run('2. COMPENSATION')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('The Client agrees to pay the Provider a total contract value of ')
    run.font.size = Pt(11)
    run = p.add_run('<ContractValue>')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(' for the services rendered under this Agreement. Payment shall be made in equal monthly installments over the term of the contract.')
    run.font.size = Pt(11)

    doc.add_paragraph('')

    # Term section
    p = doc.add_paragraph()
    run = p.add_run('3. TERM AND TERMINATION')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('This Agreement shall be effective as of ')
    run.font.size = Pt(11)
    run = p.add_run('<EffectiveDate>')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(' and shall remain in force for a period of ')
    run.font.size = Pt(11)
    run = p.add_run('<TermLength>')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(', unless terminated earlier in accordance with the provisions of this section.')
    run.font.size = Pt(11)

    doc.add_paragraph('')

    # Confidentiality section
    p = doc.add_paragraph()
    run = p.add_run('4. CONFIDENTIALITY')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('Both Parties agree to maintain the confidentiality of all proprietary information exchanged during the course of this Agreement. This obligation shall survive the termination of the Agreement for a period of two (2) years.')
    run.font.size = Pt(11)

    doc.add_paragraph('')

    # Signatures
    p = doc.add_paragraph()
    run = p.add_run('SIGNATURES')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('_________________________________')
    run.font.size = Pt(11)
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('For Global Enterprises Corp.')
    run.font.size = Pt(11)

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('_________________________________')
    run.font.size = Pt(11)
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('For <PartyName>')
    run.font.size = Pt(11)

    doc.save(TEMPLATE_FILE)
    print(f'Template created: {TEMPLATE_FILE}')


def create_data_source():
    """Create the CSV data source with 15 contract records."""
    with open(DATA_FILE, 'w', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=['PartyName', 'ContractValue', 'EffectiveDate', 'TermLength'])
        writer.writeheader()
        for rec in CONTRACT_RECORDS:
            writer.writerow(rec)
    print(f'Data source created: {DATA_FILE}')


def main():
    create_template()
    create_data_source()

    # Open template in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{TEMPLATE_FILE}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with Contract_Template.docx')


main()
