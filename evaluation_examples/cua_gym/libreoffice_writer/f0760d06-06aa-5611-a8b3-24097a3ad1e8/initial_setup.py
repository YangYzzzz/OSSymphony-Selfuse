"""
Initial Setup: Open a personal letter document in LibreOffice Writer
Task ID: writer_lec_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up standard A4 letter page (default before envelope creation)
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Sender info at the top
    sender = doc.add_paragraph()
    sender.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = sender.add_run("Anna Becker")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"
    run = sender.add_run("\nBlumenweg 12")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"
    run = sender.add_run("\n10115 Berlin, Germany")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    # Date
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_para.paragraph_format.space_before = Pt(24)
    run = date_para.add_run("15. March 2025")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    # Recipient address block (in letter body, not envelope)
    doc.add_paragraph()  # blank line
    recipient = doc.add_paragraph()
    recipient.paragraph_format.space_before = Pt(12)
    for line in ["Herr Klaus Mueller", "Hauptstrasse 45", "80331 Munchen", "Germany"]:
        run = recipient.add_run(line + "\n")
        run.font.size = Pt(11)
        run.font.name = "Liberation Serif"

    # Greeting
    doc.add_paragraph()
    greeting = doc.add_paragraph()
    run = greeting.add_run("Sehr geehrter Herr Mueller,")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    # Letter body
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_before = Pt(12)
    run = body1.add_run(
        "I hope this letter finds you well. I am writing to follow up on our "
        "recent conversation regarding the upcoming conference in Munich. As we "
        "discussed, the event is scheduled for the first week of May, and I wanted "
        "to confirm the arrangements we outlined."
    )
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_before = Pt(6)
    run = body2.add_run(
        "Please find enclosed the preliminary agenda and the list of speakers who "
        "have confirmed their participation. I would appreciate it if you could "
        "review these documents and share your thoughts at your earliest convenience."
    )
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_before = Pt(6)
    run = body3.add_run(
        "Looking forward to meeting you in Munich. Please do not hesitate to "
        "contact me if you have any questions or require further information."
    )
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    # Closing
    doc.add_paragraph()
    closing = doc.add_paragraph()
    run = closing.add_run("Mit freundlichen Gruessen,")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    doc.add_paragraph()
    sig = doc.add_paragraph()
    run = sig.add_run("Anna Becker")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
