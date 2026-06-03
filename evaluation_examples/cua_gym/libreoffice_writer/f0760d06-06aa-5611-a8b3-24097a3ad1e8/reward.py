"""
Reward Script: Create C6 envelope document for European personal letter
Task ID: writer_lec_057
Domain: libreoffice_writer
Scoring:
  Component 1: C6 page dimensions (162x114mm) — 0.30 points
  Component 2: Landscape orientation — 0.10 points
  Component 3: Delivery address content and formatting — 0.30 points
  Component 4: Return address present with smaller font — 0.15 points
  Component 5: Letter body removed (envelope only) — 0.15 points
"""

import os
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_057'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Pre-calculate page dimensions for use across components
    pw_mm = section.page_width / 36000.0
    ph_mm = section.page_height / 36000.0

    # Component 1: C6 page dimensions — 162mm x 114mm (0.30 points)
    # C6 envelope is 162mm wide x 114mm tall (or vice versa in landscape)
    # Allow ~3mm tolerance for rounding
    try:
        # Check for C6 dimensions in either orientation encoding
        is_c6 = False
        if abs(pw_mm - 162.0) < 3.0 and abs(ph_mm - 114.0) < 3.0:
            is_c6 = True
        elif abs(pw_mm - 114.0) < 3.0 and abs(ph_mm - 162.0) < 3.0:
            is_c6 = True

        if is_c6:
            print(f"PASS: Component 1 — C6 dimensions detected: {pw_mm:.1f}x{ph_mm:.1f}mm (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — Expected C6 (162x114mm), found: {pw_mm:.1f}x{ph_mm:.1f}mm")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Landscape orientation (0.10 points)
    # Envelope documents are typically landscape. The golden file has orientation=LANDSCAPE
    try:
        is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
        # Also accept if width > height (physical landscape even if orientation flag differs)
        width_gt_height = section.page_width > section.page_height
        if is_landscape or width_gt_height:
            print(f"PASS: Component 2 — Landscape orientation (orient={section.orientation}, w>h={width_gt_height}) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Expected landscape, found orientation={section.orientation}, w={section.page_width}, h={section.page_height}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Delivery address on envelope with proper formatting (0.30 points)
    # COMPOUND CHECK: Delivery address must be on a C6-sized document (not the original A4 letter)
    # AND in a larger font (>=10pt). This ensures we only score the envelope version.
    try:
        # Gate: page must NOT be A4 (must be envelope-sized, i.e. significantly smaller than A4)
        page_is_envelope = (pw_mm < 200.0 and ph_mm < 200.0)  # A4 is 210x297

        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        delivery_parts = ["klaus mueller", "hauptstrasse 45", "80331", "munchen", "germany"]
        parts_found = sum(1 for part in delivery_parts if part in all_text)

        # Check that delivery address runs are in a larger font (>=10pt)
        delivery_font_ok = False
        for p in doc.paragraphs:
            p_text = p.text.lower()
            if "klaus mueller" in p_text or "hauptstrasse 45" in p_text:
                for r in p.runs:
                    if r.font.size and r.font.size.pt >= 10.0:
                        delivery_font_ok = True
                        break
                if delivery_font_ok:
                    break

        sub_score = 0.0
        if page_is_envelope and parts_found >= 4:
            sub_score += 0.20
            print(f"PASS: Component 3a — Delivery address on envelope ({parts_found}/5 parts)")
        elif not page_is_envelope:
            print(f"FAIL: Component 3a — Page is not envelope-sized ({pw_mm:.1f}x{ph_mm:.1f}mm)")
        else:
            print(f"FAIL: Component 3a — Only {parts_found}/5 delivery address parts found")

        if page_is_envelope and delivery_font_ok:
            sub_score += 0.10
            print(f"PASS: Component 3b — Delivery address font >=10pt on envelope (0.10 pts)")
        elif not page_is_envelope:
            print(f"FAIL: Component 3b — Not on envelope-sized page")
        else:
            print(f"FAIL: Component 3b — Delivery address font not >=10pt or not found")

        if sub_score > 0:
            print(f"  Component 3 total: {sub_score:.2f} pts")
            total_score += sub_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Return address with smaller font on envelope (0.15 points)
    # COMPOUND CHECK: Return address must be on envelope-sized page AND in smaller font (<=10pt)
    # In initial A4 letter, sender address is at 11pt — this check requires both envelope + small font
    try:
        page_is_envelope = (pw_mm < 200.0 and ph_mm < 200.0)

        return_addr_found = False
        return_addr_small_font = False
        for p in doc.paragraphs:
            p_text = p.text.lower()
            if "anna becker" in p_text or "blumenweg" in p_text or "10115 berlin" in p_text:
                return_addr_found = True
                for r in p.runs:
                    if r.font.size and r.font.size.pt <= 10.0:
                        return_addr_small_font = True
                        break
                break

        if page_is_envelope and return_addr_found and return_addr_small_font:
            print(f"PASS: Component 4 — Return address on envelope with small font (0.15 pts)")
            total_score += 0.15
        elif not page_is_envelope:
            print(f"FAIL: Component 4 — Not on envelope-sized page")
        elif not return_addr_found:
            print(f"FAIL: Component 4 — Return address not found")
        else:
            print(f"FAIL: Component 4 — Return address font not <=10pt")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Letter body removed — envelope only (0.15 points)
    # The initial file has letter body paragraphs (greeting, body text, closing).
    # The golden file has only envelope content (addresses, spacers).
    # Check that typical letter-body phrases are absent.
    try:
        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        letter_phrases = [
            "sehr geehrter",
            "i hope this letter finds you well",
            "mit freundlichen gruessen",
            "looking forward to meeting",
            "preliminary agenda"
        ]
        letter_phrases_found = sum(1 for phrase in letter_phrases if phrase in full_text)

        if letter_phrases_found == 0:
            print(f"PASS: Component 5 — Letter body removed, envelope only (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — {letter_phrases_found} letter body phrases still present")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
