"""
Initial Setup: Create a technical documentation document with 5 section headings
each having a default black bottom border.
Task ID: writer_tech_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bottom_border(paragraph, color_hex="000000", size=12):
    """Add a bottom border to a paragraph using XML.
    size is in eighth-points (12 = 1.5pt).
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Document title
    title = doc.add_heading('CloudSync Platform — Technical Documentation', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This document provides comprehensive technical documentation for the '
        'CloudSync Platform v3.2, covering architecture, APIs, deployment, '
        'security, and monitoring subsystems.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # --- Section 1: Architecture Overview ---
    h1 = doc.add_heading('1. Architecture Overview', level=1)
    add_bottom_border(h1, color_hex="000000", size=12)

    doc.add_paragraph(
        'The CloudSync Platform employs a microservices architecture deployed '
        'across three availability zones in the us-east-1 region. The system '
        'processes an average of 2.4 million synchronization events daily, with '
        'peak throughput reaching 847 events per second during business hours.'
    )
    doc.add_paragraph(
        'Core services include the Sync Engine (Go 1.21), the Conflict Resolver '
        '(Rust 1.74), and the Metadata Index (Python 3.12 with FastAPI). Each '
        'service communicates via gRPC with TLS 1.3 mutual authentication. The '
        'message bus uses Apache Kafka 3.6 with a 7-day retention policy and '
        'compacted topics for configuration state.'
    )
    doc.add_paragraph(
        'Data persistence relies on PostgreSQL 16 for relational metadata, '
        'Redis 7.2 for session caching, and MinIO for object storage with '
        'erasure coding (EC:4+2) across six storage nodes.'
    )

    # --- Section 2: API Reference ---
    h2 = doc.add_heading('2. API Reference', level=1)
    add_bottom_border(h2, color_hex="000000", size=12)

    doc.add_paragraph(
        'All public-facing APIs follow the OpenAPI 3.1 specification and are '
        'served through an Envoy proxy with rate limiting configured at 1,000 '
        'requests per minute per API key. Authentication uses OAuth 2.0 with '
        'PKCE flow for browser-based clients and JWT bearer tokens for '
        'service-to-service calls.'
    )
    doc.add_paragraph(
        'The primary sync endpoint POST /api/v3/sync accepts multipart payloads '
        'up to 256 MB. Chunked uploads are supported for files exceeding 50 MB '
        'via the PUT /api/v3/upload/{session_id}/chunk/{index} endpoint. Each '
        'chunk must be exactly 5 MB except for the final chunk.'
    )
    doc.add_paragraph(
        'Webhook notifications are dispatched to registered endpoints within '
        '30 seconds of event completion. The payload includes a SHA-256 HMAC '
        'signature in the X-CloudSync-Signature header for verification. Failed '
        'deliveries are retried with exponential backoff (1s, 2s, 4s, 8s, 16s) '
        'for up to 24 hours.'
    )

    # --- Section 3: Deployment Pipeline ---
    h3 = doc.add_heading('3. Deployment Pipeline', level=1)
    add_bottom_border(h3, color_hex="000000", size=12)

    doc.add_paragraph(
        'Continuous deployment is managed through a GitOps workflow using '
        'ArgoCD 2.9 with automatic sync enabled for the staging environment '
        'and manual approval gates for production. The CI pipeline runs on '
        'GitHub Actions with self-hosted ARM64 runners for cost optimization.'
    )
    doc.add_paragraph(
        'Container images are built with Buildpacks and stored in Harbor '
        'registry with Trivy vulnerability scanning. Any CVE rated CRITICAL '
        'or HIGH blocks promotion to production. Images are signed with Cosign '
        'and verified at deployment via Kyverno admission policies.'
    )
    doc.add_paragraph(
        'Canary deployments use Flagger with progressive traffic shifting: '
        '5% initial, then 10%, 25%, 50%, 75%, and 100% over 45 minutes. '
        'Automated rollback triggers if error rate exceeds 0.1% or p99 latency '
        'exceeds 500ms during any stage.'
    )

    # --- Section 4: Security Controls ---
    h4 = doc.add_heading('4. Security Controls', level=1)
    add_bottom_border(h4, color_hex="000000", size=12)

    doc.add_paragraph(
        'Access control implements a zero-trust model with identity-aware '
        'proxy (IAP) for all internal services. Service mesh policies enforce '
        'mTLS between all pods, and network policies restrict east-west traffic '
        'to explicitly declared dependencies in the service catalog.'
    )
    doc.add_paragraph(
        'Encryption at rest uses AES-256-GCM with customer-managed keys stored '
        'in HashiCorp Vault 1.15 with auto-unseal via AWS KMS. Key rotation '
        'occurs every 90 days with a 7-day grace period for re-encryption of '
        'active sessions. All secrets are injected via the Vault Agent sidecar.'
    )
    doc.add_paragraph(
        'Security scanning runs continuously: SAST via Semgrep with custom '
        'rules, DAST via ZAP against staging endpoints, and SCA via Snyk for '
        'dependency monitoring. Findings are tracked in Jira with SLA of 24 '
        'hours for critical and 7 days for high-severity issues.'
    )

    # --- Section 5: Monitoring & Observability ---
    h5 = doc.add_heading('5. Monitoring & Observability', level=1)
    add_bottom_border(h5, color_hex="000000", size=12)

    doc.add_paragraph(
        'Observability is built on the OpenTelemetry stack with traces exported '
        'to Jaeger, metrics to Prometheus with Thanos for long-term storage, '
        'and logs to Loki with a 30-day retention policy. Grafana dashboards '
        'provide real-time visibility into sync throughput, error rates, and '
        'resource utilization across all services.'
    )
    doc.add_paragraph(
        'Alerting uses a tiered notification system: P1 incidents page the '
        'on-call engineer via PagerDuty within 60 seconds, P2 alerts go to the '
        'team Slack channel, and P3 issues create Jira tickets for the next '
        'sprint. SLOs are defined at 99.95% availability and 200ms p95 latency '
        'for the sync API.'
    )
    doc.add_paragraph(
        'Custom metrics track business KPIs including sync success rate '
        '(target: 99.99%), conflict resolution time (target: < 100ms), and '
        'data freshness (target: < 5 minutes for 95th percentile). Monthly '
        'SLO reports are generated automatically and distributed to stakeholders.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
