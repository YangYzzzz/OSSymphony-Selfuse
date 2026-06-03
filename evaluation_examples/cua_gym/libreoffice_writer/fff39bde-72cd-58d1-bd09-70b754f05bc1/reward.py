"""
Reward Script: Employee Satisfaction Survey Results Report
Task ID: writer_hr_076
Domain: libreoffice_writer
Scoring:
  Component 1: Tables exist (>= 8 tables) — 0.15
  Component 2: Heading structure (Heading 1/2 styles) — 0.15
  Component 3: Survey question tables (5 category tables with scores) — 0.15
  Component 4: Conditional formatting on score cells (green/yellow/red) — 0.20
  Component 5: Benchmark comparison table (>= 5 metrics) — 0.10
  Component 6: Demographic breakout tables (3 dimensions) — 0.10
  Component 7: Action plan table with required columns — 0.15
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_076'


def get_cell_fill(cell):
    """Extract background fill color from a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill.lower() != 'auto':
                return fill.upper()
    return None


def is_green_fill(fill):
    """Check if fill color is a shade of green."""
    if not fill:
        return False
    # Common greens: 00B050, 00FF00, 008000, 92D050, etc.
    try:
        r = int(fill[0:2], 16)
        g = int(fill[2:4], 16)
        b = int(fill[4:6], 16)
        return g > 100 and g > r and g > b
    except (ValueError, IndexError):
        return False


def is_yellow_fill(fill):
    """Check if fill color is a shade of yellow/amber."""
    if not fill:
        return False
    # Common yellows: FFC000, FFFF00, FFD700, etc.
    try:
        r = int(fill[0:2], 16)
        g = int(fill[2:4], 16)
        b = int(fill[4:6], 16)
        return r > 150 and g > 100 and b < 100 and r >= g
    except (ValueError, IndexError):
        return False


def is_red_fill(fill):
    """Check if fill color is a shade of red."""
    if not fill:
        return False
    # Common reds: FF0000, FF4444, CC0000, etc.
    try:
        r = int(fill[0:2], 16)
        g = int(fill[2:4], 16)
        b = int(fill[4:6], 16)
        return r > 150 and g < 80 and b < 80
    except (ValueError, IndexError):
        return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_tables = len(doc.tables)
    num_paras = len(doc.paragraphs)

    # Component 1: Tables exist — at least 8 tables (0.15 points)
    # Initial has 0 tables, golden has 12
    try:
        if num_tables >= 8:
            print(f"PASS: Component 1 — Document has {num_tables} tables (>= 8 required) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Document has {num_tables} tables, need >= 8")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Heading structure with Heading 1 and Heading 2 styles (0.15 points)
    # Initial has all Normal style paragraphs, golden uses Heading 1/2
    try:
        heading1_count = 0
        heading2_count = 0
        for para in doc.paragraphs:
            if para.style and para.style.name == 'Heading 1':
                heading1_count += 1
            elif para.style and para.style.name == 'Heading 2':
                heading2_count += 1

        if heading1_count >= 4 and heading2_count >= 3:
            print(f"PASS: Component 2 — Found {heading1_count} Heading 1 and {heading2_count} Heading 2 styles (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Heading 1 count={heading1_count} (need >= 4), Heading 2 count={heading2_count} (need >= 3)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Survey question tables — 5 category tables with Q# and score columns (0.15 points)
    # Each category table should have Question # header and score data rows
    try:
        survey_table_count = 0
        for table in doc.tables:
            rows = table.rows
            if len(rows) >= 4:  # header + at least 3 questions
                header_cells = [c.text.strip().lower() for c in rows[0].cells]
                header_text = ' '.join(header_cells)
                # Check for question/score column pattern
                has_question_col = any('question' in h or 'q#' in h for h in header_cells)
                has_score_col = any('score' in h or 'avg' in h for h in header_cells)
                if has_question_col and has_score_col:
                    # Verify data rows have Q-numbered entries
                    first_data = rows[1].cells[0].text.strip()
                    if first_data.startswith('Q'):
                        survey_table_count += 1

        if survey_table_count >= 5:
            print(f"PASS: Component 3 — Found {survey_table_count} survey question tables (>= 5) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — Found {survey_table_count} survey question tables, need >= 5")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Conditional formatting — score cells have green/yellow/red backgrounds (0.20 points)
    # Task requires: green (4.0+), yellow (3.0-3.9), red (below 3.0)
    # Initial has no tables so no cell colors; golden has color-coded score cells
    try:
        colored_correct = 0
        colored_total = 0
        has_green = False
        has_yellow = False
        has_red = False

        for table in doc.tables:
            rows = table.rows
            if len(rows) < 4:
                continue
            header_cells = [c.text.strip().lower() for c in rows[0].cells]
            has_question_col = any('question' in h or 'q#' in h for h in header_cells)
            has_score_col = any('score' in h or 'avg' in h for h in header_cells)
            if not (has_question_col and has_score_col):
                continue

            # Find score column index
            score_col = -1
            for ci, h in enumerate(header_cells):
                if 'score' in h or 'avg' in h:
                    score_col = ci
                    break

            if score_col < 0:
                continue

            for ri in range(1, len(rows)):
                cell = rows[ri].cells[score_col]
                cell_text = cell.text.strip()
                try:
                    score_val = float(cell_text)
                except (ValueError, TypeError):
                    continue

                fill = get_cell_fill(cell)
                colored_total += 1

                if score_val >= 4.0 and is_green_fill(fill):
                    colored_correct += 1
                    has_green = True
                elif 3.0 <= score_val < 4.0 and is_yellow_fill(fill):
                    colored_correct += 1
                    has_yellow = True
                elif score_val < 3.0 and is_red_fill(fill):
                    colored_correct += 1
                    has_red = True
                elif fill is not None:
                    # Has some color, check if it's at least in the right category
                    pass

        if colored_total >= 10 and has_green and has_yellow and has_red:
            color_ratio = colored_correct / colored_total
            if color_ratio >= 0.8:
                print(f"PASS: Component 4 — {colored_correct}/{colored_total} cells correctly colored, all 3 colors present (0.20 pts)")
                total_score += 0.20
            else:
                partial = round(0.20 * color_ratio, 2)
                print(f"PARTIAL: Component 4 — {colored_correct}/{colored_total} cells correct ({partial} pts)")
                total_score += partial
        elif colored_total > 0 and (has_green or has_yellow or has_red):
            partial = 0.10
            print(f"PARTIAL: Component 4 — Some color coding found: {colored_correct}/{colored_total} correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No conditional formatting found (colored_total={colored_total})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Benchmark comparison table — at least 5 metrics with company/industry columns (0.10 points)
    try:
        benchmark_found = False
        for table in doc.tables:
            rows = table.rows
            if len(rows) < 6:  # header + at least 5 metrics
                continue
            header_cells = [c.text.strip().lower() for c in rows[0].cells]
            header_text = ' '.join(header_cells)
            has_metric = any('metric' in h for h in header_cells)
            has_company = any('company' in h for h in header_cells)
            has_industry = any('industry' in h for h in header_cells)

            if has_metric and has_company and has_industry:
                benchmark_found = True
                print(f"PASS: Component 5 — Benchmark table found with {len(rows)-1} metrics (0.10 pts)")
                total_score += 0.10
                break

        if not benchmark_found:
            print(f"FAIL: Component 5 — No benchmark comparison table found with Metric/Company/Industry columns")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Demographic breakout tables — 3 dimensions (dept/tenure/level) (0.10 points)
    # Each should have Work Env, Mgmt, Career, Comp, Culture columns
    try:
        demo_tables = 0
        for table in doc.tables:
            rows = table.rows
            if len(rows) < 4:
                continue
            header_cells = [c.text.strip().lower() for c in rows[0].cells]

            # Check first column header for demographic dimension
            first_col = header_cells[0] if header_cells else ''
            is_dept = 'department' in first_col or 'dept' in first_col
            is_tenure = 'tenure' in first_col
            is_level = 'level' in first_col or 'job' in first_col

            if is_dept or is_tenure or is_level:
                # Check for category score columns (at least 3 of: work env, mgmt, career, comp, culture)
                cat_cols = 0
                for h in header_cells[1:]:
                    if any(kw in h for kw in ['work', 'mgmt', 'management', 'career', 'comp', 'culture']):
                        cat_cols += 1
                if cat_cols >= 3:
                    demo_tables += 1

        if demo_tables >= 3:
            print(f"PASS: Component 6 — Found {demo_tables} demographic breakout tables (0.10 pts)")
            total_score += 0.10
        elif demo_tables >= 1:
            partial = round(0.10 * demo_tables / 3, 2)
            print(f"PARTIAL: Component 6 — Found {demo_tables}/3 demographic breakout tables ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — No demographic breakout tables found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Action plan table with required columns (0.15 points)
    # Must have: Focus Area, Current Score, Target, Owner, Timeline
    try:
        action_plan_found = False
        for table in doc.tables:
            rows = table.rows
            if len(rows) < 3:  # header + at least 2 action items
                continue
            header_cells = [c.text.strip().lower() for c in rows[0].cells]

            has_focus = any('focus' in h or 'area' in h or 'initiative' in h for h in header_cells)
            has_current = any('current' in h for h in header_cells)
            has_target = any('target' in h for h in header_cells)
            has_owner = any('owner' in h for h in header_cells)
            has_timeline = any('timeline' in h or 'deadline' in h or 'by' in h.split() for h in header_cells)

            matched = sum([has_focus, has_current, has_target, has_owner, has_timeline])

            if matched >= 4:
                action_plan_found = True
                print(f"PASS: Component 7 — Action plan table found with {matched}/5 required columns, {len(rows)-1} items (0.15 pts)")
                total_score += 0.15
                break

        if not action_plan_found:
            print(f"FAIL: Component 7 — No action plan table with Focus Area/Current Score/Target/Owner/Timeline")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
