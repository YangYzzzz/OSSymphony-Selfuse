"""
Initial Setup: Employee satisfaction survey results - raw data document
Task ID: writer_hr_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title - just plain text, not a structured report
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Nextera Solutions Inc. - Employee Satisfaction Survey 2025')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('')

    # Raw narrative intro
    doc.add_paragraph(
        'Below is the raw data collected from the 2025 Annual Employee Satisfaction Survey '
        'conducted between February 3 and February 28, 2025. The survey was distributed to '
        'all 847 full-time employees across six departments. A total of 694 responses were '
        'received, yielding an overall response rate of 81.9%. The survey used a 5-point '
        'Likert scale (1 = Strongly Disagree, 5 = Strongly Agree).'
    )

    doc.add_paragraph('')
    doc.add_paragraph('--- RAW SURVEY DATA ---')
    doc.add_paragraph('')

    # Raw response rate data as plain text
    doc.add_paragraph('RESPONSE RATES BY DEPARTMENT:')
    dept_rates = [
        ('Engineering', 142, 168, '84.5%'),
        ('Marketing', 89, 104, '85.6%'),
        ('Sales', 118, 152, '77.6%'),
        ('Human Resources', 52, 58, '89.7%'),
        ('Finance', 78, 91, '85.7%'),
        ('Operations', 215, 274, '78.5%'),
    ]
    for dept, resp, total, rate in dept_rates:
        doc.add_paragraph(f'  {dept}: {resp} out of {total} ({rate})')

    doc.add_paragraph('')
    doc.add_paragraph('SURVEY QUESTIONS AND AVERAGE SCORES:')
    doc.add_paragraph('')

    # Category 1: Work Environment
    doc.add_paragraph('Category 1 - Work Environment')
    work_env_qs = [
        ('Q1', 'My workspace is comfortable and well-equipped', 4.2),
        ('Q2', 'The office facilities meet my needs', 3.8),
        ('Q3', 'I feel safe in my work environment', 4.5),
        ('Q4', 'The noise level in my workspace is acceptable', 3.1),
        ('Q5', 'I have access to necessary tools and technology', 4.0),
    ]
    for qid, text, score in work_env_qs:
        doc.add_paragraph(f'  {qid}. {text} -- Avg: {score}')

    doc.add_paragraph('')

    # Category 2: Management & Leadership
    doc.add_paragraph('Category 2 - Management & Leadership')
    mgmt_qs = [
        ('Q6', 'My direct manager provides clear expectations', 3.9),
        ('Q7', 'Senior leadership communicates company vision effectively', 3.2),
        ('Q8', 'I receive regular and constructive feedback', 2.8),
        ('Q9', 'Management is open to employee suggestions', 3.5),
        ('Q10', 'Decisions are made transparently', 2.9),
    ]
    for qid, text, score in mgmt_qs:
        doc.add_paragraph(f'  {qid}. {text} -- Avg: {score}')

    doc.add_paragraph('')

    # Category 3: Career Development
    doc.add_paragraph('Category 3 - Career Development')
    career_qs = [
        ('Q11', 'I have opportunities for professional growth', 3.4),
        ('Q12', 'Training programs are relevant and accessible', 3.0),
        ('Q13', 'I see a clear career path at this company', 2.7),
        ('Q14', 'My skills are being utilized effectively', 3.6),
        ('Q15', 'Promotion criteria are fair and transparent', 2.5),
    ]
    for qid, text, score in career_qs:
        doc.add_paragraph(f'  {qid}. {text} -- Avg: {score}')

    doc.add_paragraph('')

    # Category 4: Compensation & Benefits
    doc.add_paragraph('Category 4 - Compensation & Benefits')
    comp_qs = [
        ('Q16', 'My salary is competitive for my role', 3.3),
        ('Q17', 'The benefits package meets my needs', 4.1),
        ('Q18', 'The bonus structure is fair', 2.6),
        ('Q19', 'I am satisfied with paid time off policies', 4.3),
        ('Q20', 'Retirement plan options are adequate', 3.7),
    ]
    for qid, text, score in comp_qs:
        doc.add_paragraph(f'  {qid}. {text} -- Avg: {score}')

    doc.add_paragraph('')

    # Category 5: Culture & Engagement
    doc.add_paragraph('Category 5 - Culture & Engagement')
    culture_qs = [
        ('Q21', 'I feel a sense of belonging at this company', 3.8),
        ('Q22', 'Collaboration across teams is encouraged', 4.0),
        ('Q23', 'The company values diversity and inclusion', 4.2),
        ('Q24', 'I would recommend this company as a place to work', 3.5),
        ('Q25', 'I feel motivated to go above and beyond', 3.1),
    ]
    for qid, text, score in culture_qs:
        doc.add_paragraph(f'  {qid}. {text} -- Avg: {score}')

    doc.add_paragraph('')

    # Raw benchmark data
    doc.add_paragraph('BENCHMARK DATA (Company vs Industry Average):')
    benchmarks = [
        ('Overall Satisfaction', 3.52, 3.65),
        ('Employee Engagement', 3.58, 3.70),
        ('Management Effectiveness', 3.26, 3.45),
        ('Career Development', 3.04, 3.30),
        ('Compensation Satisfaction', 3.60, 3.55),
        ('Work-Life Balance', 3.92, 3.80),
        ('Culture & Belonging', 3.72, 3.60),
    ]
    for metric, company, industry in benchmarks:
        doc.add_paragraph(f'  {metric}: Company={company}, Industry={industry}')

    doc.add_paragraph('')

    # Demographic breakout raw data
    doc.add_paragraph('DEMOGRAPHIC BREAKOUT - Average Scores:')
    doc.add_paragraph('')

    doc.add_paragraph('By Department:')
    dept_scores = [
        ('Engineering', 3.8),
        ('Marketing', 3.6),
        ('Sales', 3.2),
        ('Human Resources', 4.0),
        ('Finance', 3.5),
        ('Operations', 3.1),
    ]
    for dept, score in dept_scores:
        doc.add_paragraph(f'  {dept}: {score}')

    doc.add_paragraph('')
    doc.add_paragraph('By Tenure:')
    tenure_scores = [
        ('Less than 1 year', 3.9),
        ('1-3 years', 3.5),
        ('3-5 years', 3.2),
        ('5-10 years', 3.0),
        ('10+ years', 3.7),
    ]
    for tenure, score in tenure_scores:
        doc.add_paragraph(f'  {tenure}: {score}')

    doc.add_paragraph('')
    doc.add_paragraph('By Job Level:')
    level_scores = [
        ('Entry Level', 3.4),
        ('Mid-Level', 3.3),
        ('Senior', 3.6),
        ('Manager', 3.8),
        ('Director+', 4.1),
    ]
    for level, score in level_scores:
        doc.add_paragraph(f'  {level}: {score}')

    doc.add_paragraph('')

    # Open-ended response raw themes
    doc.add_paragraph('OPEN-ENDED RESPONSE THEMES (Top mentions):')
    themes = [
        ('Better communication from leadership', 187),
        ('More career development opportunities', 156),
        ('Improved work-life balance', 134),
        ('Updated office facilities/equipment', 98),
        ('More competitive compensation', 89),
        ('Better cross-team collaboration', 76),
        ('Enhanced remote work flexibility', 72),
        ('Recognition and appreciation programs', 65),
    ]
    for theme, count in themes:
        doc.add_paragraph(f'  - {theme} ({count} mentions)')

    doc.add_paragraph('')

    # Raw action items notes
    doc.add_paragraph('PROPOSED ACTION ITEMS (Notes):')
    actions = [
        ('Leadership Communication', 2.9, 3.5, 'VP People Ops', 'Q3 2025'),
        ('Career Pathing Program', 2.7, 3.3, 'HR Director', 'Q2 2025'),
        ('Feedback Culture Initiative', 2.8, 3.5, 'CHRO', 'Q3 2025'),
        ('Compensation Review', 2.6, 3.2, 'CFO', 'Q4 2025'),
        ('Office Modernization', 3.1, 4.0, 'Facilities Manager', 'Q1 2026'),
    ]
    for area, current, target, owner, timeline in actions:
        doc.add_paragraph(f'  {area} | Current: {current} | Target: {target} | Owner: {owner} | By: {timeline}')

    doc.add_paragraph('')
    doc.add_paragraph('--- END OF RAW DATA ---')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
