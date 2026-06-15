"""
Initial Setup: Report cover page formatting task
Task ID: writer_txtfmt_080
Domain: libreoffice_writer
Creates final_report.docx on the Desktop with a cover page and report body.
All cover page text is in 12pt Times New Roman regular black (unformatted).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'final_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Cover Page ---
    # Line 1: 'FINAL REPORT' - 12pt Times New Roman, regular, black
    para1 = doc.add_paragraph()
    run1 = para1.add_run('FINAL REPORT')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    run1.bold = False
    run1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Line 2: Subtitle - 12pt Times New Roman, regular, black
    para2 = doc.add_paragraph()
    run2 = para2.add_run('Comprehensive Market Analysis for Q1-Q4 2024')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.bold = False
    run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Line 3: Author line - 12pt Times New Roman, regular, black
    para3 = doc.add_paragraph()
    run3 = para3.add_run('Prepared by: Strategic Planning Division')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    run3.bold = False
    run3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # --- Report Body ---
    doc.add_paragraph()  # blank line separator

    # Executive Summary section
    heading_para = doc.add_paragraph()
    h_run = heading_para.add_run('Executive Summary')
    h_run.font.name = 'Times New Roman'
    h_run.font.size = Pt(14)
    h_run.bold = True

    body1 = doc.add_paragraph()
    b1_run = body1.add_run(
        'This comprehensive market analysis covers the performance metrics and strategic developments '
        'observed across the full fiscal year 2024. The report aggregates data from all four quarters, '
        'providing insights into revenue trends, market share dynamics, and competitive positioning '
        'across key industry segments.'
    )
    b1_run.font.name = 'Times New Roman'
    b1_run.font.size = Pt(11)

    body2 = doc.add_paragraph()
    b2_run = body2.add_run(
        'Total revenue for fiscal year 2024 reached $4.82 billion, representing a 12.4% increase '
        'compared to the previous year. The strongest growth was observed in the Asia-Pacific region, '
        'where new market entries and expanded distribution channels contributed to a 23.7% uplift '
        'in sales volume during Q3 and Q4.'
    )
    b2_run.font.name = 'Times New Roman'
    b2_run.font.size = Pt(11)

    # Market Performance section
    perf_heading = doc.add_paragraph()
    ph_run = perf_heading.add_run('Market Performance Overview')
    ph_run.font.name = 'Times New Roman'
    ph_run.font.size = Pt(14)
    ph_run.bold = True

    perf_body = doc.add_paragraph()
    pb_run = perf_body.add_run(
        'Q1 2024 established a strong foundation with $1.12 billion in net revenues, driven primarily '
        'by the successful launch of the Enterprise Solutions Suite and expanded government contracts '
        'in North America. The strategic planning division coordinated cross-functional initiatives '
        'that resulted in a 15% reduction in operational overhead.'
    )
    pb_run.font.name = 'Times New Roman'
    pb_run.font.size = Pt(11)

    perf_body2 = doc.add_paragraph()
    pb2_run = perf_body2.add_run(
        'Q2 2024 showed moderate growth at $1.18 billion despite headwinds from supply chain '
        'disruptions affecting the semiconductor and logistics sectors. The strategic pivot toward '
        'cloud-based service offerings began to yield measurable returns, with subscription revenues '
        'growing 31.2% year-over-year.'
    )
    pb2_run.font.name = 'Times New Roman'
    pb2_run.font.size = Pt(11)

    perf_body3 = doc.add_paragraph()
    pb3_run = perf_body3.add_run(
        'Q3 2024 marked the strongest quarter at $1.29 billion, bolstered by the acquisition of '
        'Meridian Analytics Corp and the expansion of our artificial intelligence product line. '
        'Customer retention rates improved to 94.3%, reflecting successful customer success '
        'program implementations across all major accounts.'
    )
    pb3_run.font.name = 'Times New Roman'
    pb3_run.font.size = Pt(11)

    perf_body4 = doc.add_paragraph()
    pb4_run = perf_body4.add_run(
        'Q4 2024 closed with $1.23 billion in revenues, reflecting seasonal normalization and '
        'deliberate investment in R&D infrastructure for the upcoming fiscal year. The strategic '
        'planning division has identified six priority growth initiatives for 2025, with projected '
        'combined impact of $340 million in incremental revenue.'
    )
    pb4_run.font.name = 'Times New Roman'
    pb4_run.font.size = Pt(11)

    # Strategic Recommendations section
    rec_heading = doc.add_paragraph()
    rh_run = rec_heading.add_run('Strategic Recommendations')
    rh_run.font.name = 'Times New Roman'
    rh_run.font.size = Pt(14)
    rh_run.bold = True

    for rec in [
        'Continue investment in AI-powered automation solutions to maintain competitive differentiation.',
        'Expand Asia-Pacific operations through strategic partnerships with regional distributors.',
        'Accelerate cloud migration initiatives to improve gross margin profile toward 68% target.',
        'Enhance supply chain resilience through dual-sourcing strategies and inventory buffer programs.',
        'Invest in talent acquisition and retention programs targeting senior engineering and data science roles.',
    ]:
        rec_para = doc.add_paragraph(style='List Bullet')
        rec_run = rec_para.add_run(rec)
        rec_run.font.name = 'Times New Roman'
        rec_run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
