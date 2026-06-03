"""
Reward Script: Comprehensive cover page formatting for final_report.docx
Task ID: writer_txtfmt_080
Domain: libreoffice_writer
Scoring:
  Component 1: FINAL REPORT — font Arial, 28pt, bold, color #008080   (0.25 pts)
  Component 2: FINAL REPORT — shadow effect present                    (0.15 pts)
  Component 3: FINAL REPORT — character spacing expanded by 4pt (80 twips)  (0.10 pts)
  Component 4: Subtitle — Georgia 16pt italic color #4A4A4A            (0.25 pts)
  Component 5: Author line — Liberation Sans 12pt SmallCaps            (0.25 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_080'
FILE_PATH = f'{WORKDIR}/Desktop/final_report.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_rpr_element(run, tag):
    """Return the named rPr child element or None."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    return rPr.find(qn(tag))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Ensure the document has at least 3 paragraphs
    if len(doc.paragraphs) < 3:
        print("CRITICAL: Document has fewer than 3 paragraphs — invalid state")
        print("REWARD: 0.0")
        return 0.0

    para0 = doc.paragraphs[0]  # FINAL REPORT
    para1 = doc.paragraphs[1]  # Subtitle
    para2 = doc.paragraphs[2]  # Author line

    # Verify paragraph texts are intact (precondition gate)
    if 'FINAL REPORT' not in para0.text:
        print("CRITICAL: Para 0 does not contain 'FINAL REPORT' — wrong paragraph structure")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: FINAL REPORT — font Arial, 28pt, bold, color #008080
    # (0.25 points)
    # -----------------------------------------------------------------------
    try:
        runs0 = [r for r in para0.runs if r.text.strip()]
        if not runs0:
            print("FAIL: Component 1 — No text runs found in para 0")
        else:
            run0 = runs0[0]
            font0 = run0.font

            # font name check
            font_name = font0.name
            size_pt = font0.size.pt if font0.size else None
            bold = font0.bold
            color = font0.color.rgb if font0.color and font0.color.type else None

            font_ok = font_name and 'Arial' in font_name
            size_ok = size_pt == 28.0
            bold_ok = bold is True
            color_ok = (color is not None and str(color).upper() == '008080')

            if font_ok and size_ok and bold_ok and color_ok:
                print(f"PASS: Component 1 — Arial 28pt Bold #008080 confirmed (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — font={font_name} (want Arial), size={size_pt} (want 28.0), "
                      f"bold={bold} (want True), color={color} (want 008080)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: FINAL REPORT — shadow effect present
    # In OOXML: <w:shadow w:val="1"/> in rPr
    # (0.15 points)
    # -----------------------------------------------------------------------
    try:
        runs0 = [r for r in para0.runs if r.text.strip()]
        if not runs0:
            print("FAIL: Component 2 — No text runs found in para 0")
        else:
            run0 = runs0[0]
            rPr = run0._element.find(qn('w:rPr'))
            shadow_el = rPr.find(qn('w:shadow')) if rPr is not None else None

            # Shadow is present if the element exists; w:val="1" or absence of val="0"
            shadow_present = False
            if shadow_el is not None:
                val = shadow_el.get(qn('w:val'))
                shadow_present = (val is None or val not in ('0', 'false', 'off'))

            if shadow_present:
                print("PASS: Component 2 — Shadow effect present on 'FINAL REPORT' (0.15 pts)")
                total_score += 0.15
            else:
                print("FAIL: Component 2 — Shadow effect NOT found on 'FINAL REPORT'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: FINAL REPORT — character spacing expanded by 4pt (80 twips)
    # In OOXML: <w:spacing w:val="80"/> in rPr (1pt = 20 twips)
    # (0.10 points)
    # -----------------------------------------------------------------------
    try:
        runs0 = [r for r in para0.runs if r.text.strip()]
        if not runs0:
            print("FAIL: Component 3 — No text runs found in para 0")
        else:
            run0 = runs0[0]
            rPr = run0._element.find(qn('w:rPr'))
            spacing_el = rPr.find(qn('w:spacing')) if rPr is not None else None

            spacing_val = None
            if spacing_el is not None:
                val_str = spacing_el.get(qn('w:val'))
                if val_str is not None:
                    spacing_val = int(val_str)

            # 4pt expanded = 4 * 20 = 80 twips
            spacing_ok = (spacing_val == 80)
            if spacing_ok:
                print(f"PASS: Component 3 — Character spacing = 80 twips (4pt expanded) (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 3 — Expected spacing=80, found spacing={spacing_val}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Subtitle — Georgia 16pt Italic color #4A4A4A
    # (0.25 points)
    # -----------------------------------------------------------------------
    try:
        runs1 = [r for r in para1.runs if r.text.strip()]
        if not runs1:
            print("FAIL: Component 4 — No text runs found in para 1 (subtitle)")
        else:
            run1 = runs1[0]
            font1 = run1.font

            font_name1 = font1.name
            size_pt1 = font1.size.pt if font1.size else None
            italic1 = font1.italic
            color1 = font1.color.rgb if font1.color and font1.color.type else None

            font_ok1 = font_name1 and 'Georgia' in font_name1
            size_ok1 = size_pt1 == 16.0
            italic_ok1 = italic1 is True
            color_ok1 = (color1 is not None and str(color1).upper() == '4A4A4A')

            if font_ok1 and size_ok1 and italic_ok1 and color_ok1:
                print(f"PASS: Component 4 — Georgia 16pt Italic #4A4A4A confirmed (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 4 — font={font_name1} (want Georgia), size={size_pt1} (want 16.0), "
                      f"italic={italic1} (want True), color={color1} (want 4A4A4A)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    # Component 5: Author line — Liberation Sans 12pt SmallCaps
    # In OOXML: <w:smallCaps w:val="1"/> in rPr
    # (0.25 points)
    # -----------------------------------------------------------------------
    try:
        runs2 = [r for r in para2.runs if r.text.strip()]
        if not runs2:
            print("FAIL: Component 5 — No text runs found in para 2 (author line)")
        else:
            run2 = runs2[0]
            font2 = run2.font

            font_name2 = font2.name
            size_pt2 = font2.size.pt if font2.size else None

            # Check SmallCaps via XML (python-docx doesn't expose smallCaps as a property)
            rPr2 = run2._element.find(qn('w:rPr'))
            small_caps_el = rPr2.find(qn('w:smallCaps')) if rPr2 is not None else None

            small_caps_on = False
            if small_caps_el is not None:
                val = small_caps_el.get(qn('w:val'))
                small_caps_on = (val is None or val not in ('0', 'false', 'off'))

            font_ok2 = font_name2 and 'Liberation Sans' in font_name2
            size_ok2 = size_pt2 == 12.0
            caps_ok = small_caps_on

            if font_ok2 and size_ok2 and caps_ok:
                print(f"PASS: Component 5 — Liberation Sans 12pt SmallCaps confirmed (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 5 — font={font_name2} (want Liberation Sans), "
                      f"size={size_pt2} (want 12.0), smallCaps={small_caps_on} (want True)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
