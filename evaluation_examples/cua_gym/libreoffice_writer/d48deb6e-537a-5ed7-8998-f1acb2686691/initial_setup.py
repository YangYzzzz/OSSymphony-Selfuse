"""
Initial Setup: Two-level bulleted list with round bullets on both levels
Task ID: writer_lec_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.oxml.ns import qn, nsmap
from lxml import etree
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Word namespace
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_numbering_definition(doc):
    """
    Add a custom numbering definition with round bullets (U+2022) on both
    level 0 and level 1.
    Returns the numId to assign to paragraphs.
    """
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part.numbering_definitions._numbering

    # Create abstractNum with two levels, both using round bullet
    abstract_num_id = '100'
    abstract_num = etree.SubElement(numbering_elm, qn('w:abstractNum'))
    abstract_num.set(qn('w:abstractNumId'), abstract_num_id)

    # Multi-level type
    multi = etree.SubElement(abstract_num, qn('w:multiLevelType'))
    multi.set(qn('w:val'), 'hybridMultilevel')

    # Level 0: round bullet U+2022
    lvl0 = etree.SubElement(abstract_num, qn('w:lvl'))
    lvl0.set(qn('w:ilvl'), '0')
    start0 = etree.SubElement(lvl0, qn('w:start'))
    start0.set(qn('w:val'), '1')
    numFmt0 = etree.SubElement(lvl0, qn('w:numFmt'))
    numFmt0.set(qn('w:val'), 'bullet')
    lvlText0 = etree.SubElement(lvl0, qn('w:lvlText'))
    lvlText0.set(qn('w:val'), '\u2022')
    lvlJc0 = etree.SubElement(lvl0, qn('w:lvlJc'))
    lvlJc0.set(qn('w:val'), 'left')
    # Indentation for level 0
    pPr0 = etree.SubElement(lvl0, qn('w:pPr'))
    ind0 = etree.SubElement(pPr0, qn('w:ind'))
    ind0.set(qn('w:left'), '720')
    ind0.set(qn('w:hanging'), '360')
    # Font for bullet
    rPr0 = etree.SubElement(lvl0, qn('w:rPr'))
    rFonts0 = etree.SubElement(rPr0, qn('w:rFonts'))
    rFonts0.set(qn('w:ascii'), 'Symbol')
    rFonts0.set(qn('w:hAnsi'), 'Symbol')
    rFonts0.set(qn('w:hint'), 'default')

    # Level 1: also round bullet U+2022
    lvl1 = etree.SubElement(abstract_num, qn('w:lvl'))
    lvl1.set(qn('w:ilvl'), '1')
    start1 = etree.SubElement(lvl1, qn('w:start'))
    start1.set(qn('w:val'), '1')
    numFmt1 = etree.SubElement(lvl1, qn('w:numFmt'))
    numFmt1.set(qn('w:val'), 'bullet')
    lvlText1 = etree.SubElement(lvl1, qn('w:lvlText'))
    lvlText1.set(qn('w:val'), '\u2022')
    lvlJc1 = etree.SubElement(lvl1, qn('w:lvlJc'))
    lvlJc1.set(qn('w:val'), 'left')
    # Indentation for level 1
    pPr1 = etree.SubElement(lvl1, qn('w:pPr'))
    ind1 = etree.SubElement(pPr1, qn('w:ind'))
    ind1.set(qn('w:left'), '1440')
    ind1.set(qn('w:hanging'), '360')
    # Font for bullet
    rPr1 = etree.SubElement(lvl1, qn('w:rPr'))
    rFonts1 = etree.SubElement(rPr1, qn('w:rFonts'))
    rFonts1.set(qn('w:ascii'), 'Symbol')
    rFonts1.set(qn('w:hAnsi'), 'Symbol')
    rFonts1.set(qn('w:hint'), 'default')

    # Insert abstractNum BEFORE any existing <w:num> elements
    first_num = numbering_elm.find(qn('w:num'))
    if first_num is not None:
        numbering_elm.insert(list(numbering_elm).index(first_num), abstract_num)

    # Create num element referencing our abstractNum
    num_id = '100'
    num_elm = etree.SubElement(numbering_elm, qn('w:num'))
    num_elm.set(qn('w:numId'), num_id)
    abstractNumId_ref = etree.SubElement(num_elm, qn('w:abstractNumId'))
    abstractNumId_ref.set(qn('w:val'), abstract_num_id)

    return int(num_id)


def add_bullet_paragraph(doc, text, level, num_id):
    """Add a paragraph as a bullet item at the given level using our numbering."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Set numbering properties
    pPr = para._element.get_or_add_pPr()
    numPr = etree.SubElement(pPr, qn('w:numPr'))
    ilvl = etree.SubElement(numPr, qn('w:ilvl'))
    ilvl.set(qn('w:val'), str(level))
    numId_elm = etree.SubElement(numPr, qn('w:numId'))
    numId_elm.set(qn('w:val'), str(num_id))

    return para


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add a title
    heading = doc.add_heading('Project Status Update', level=1)

    # Add an intro paragraph
    intro = doc.add_paragraph(
        'The following items summarize the key activities and deliverables '
        'for Q2 2025. Each main category includes specific action items '
        'that require attention from the respective teams.'
    )
    intro.runs[0].font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # We need to ensure the numbering part exists before adding our definition.
    # Adding a dummy list bullet paragraph forces python-docx to create it.
    dummy = doc.add_paragraph('', style='List Bullet')
    # Remove the dummy
    dummy._element.getparent().remove(dummy._element)

    num_id = add_numbering_definition(doc)

    # Level-1 items (4 items) and Level-2 sub-items (7 total)
    # Item 1 with 2 sub-items
    add_bullet_paragraph(doc, 'Finalize the quarterly revenue report for stakeholder review', 0, num_id)
    add_bullet_paragraph(doc, 'Consolidate regional sales data from all branch offices', 1, num_id)
    add_bullet_paragraph(doc, 'Verify tax calculations against updated fiscal guidelines', 1, num_id)

    # Item 2 with 2 sub-items
    add_bullet_paragraph(doc, 'Prepare the employee onboarding materials for June cohort', 0, num_id)
    add_bullet_paragraph(doc, 'Update the benefits enrollment handbook with 2025 plan changes', 1, num_id)
    add_bullet_paragraph(doc, 'Schedule orientation sessions with department leads', 1, num_id)

    # Item 3 with 2 sub-items
    add_bullet_paragraph(doc, 'Migrate the staging environment to the new cloud infrastructure', 0, num_id)
    add_bullet_paragraph(doc, 'Run performance benchmarks on the provisioned database cluster', 1, num_id)
    add_bullet_paragraph(doc, 'Configure automated backup schedules for critical data stores', 1, num_id)

    # Item 4 with 1 sub-item
    add_bullet_paragraph(doc, 'Launch the customer feedback survey for the spring product release', 0, num_id)
    add_bullet_paragraph(doc, 'Draft follow-up communication templates for survey respondents', 1, num_id)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
