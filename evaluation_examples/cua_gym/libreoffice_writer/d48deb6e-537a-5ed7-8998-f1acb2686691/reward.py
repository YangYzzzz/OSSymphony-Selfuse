"""
Reward Script: Change level-2 sub-items to square bullets
Task ID: writer_lec_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Level-1 bullets changed to square character (U+25A0 or U+25AA)
  Component 2 (0.3): Level-0 bullets remain round AND level-1 is square (compound)
  Component 3 (0.2): Text content/hierarchy preserved AND level-1 is square (compound)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_030'

# Square bullet characters that are acceptable
SQUARE_BULLETS = {'\u25A0', '\u25AA'}  # Black square, small black square

# Round bullet characters (including Symbol font private use area)
ROUND_BULLETS = {'\u2022', '\uF0B7'}  # Bullet, Symbol font bullet


def get_numbering_info(doc):
    """Extract bullet characters for each level from the numbering definition used by the list."""
    from docx.oxml.ns import qn

    numbering_part = doc.part.numbering_part
    if not numbering_part:
        return None, None

    numbering_xml = numbering_part._element

    # Find which abstractNumId is used by the list paragraphs
    # We look at the first list paragraph to get numId
    target_numId = None
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                numId_elem = numPr.find(qn('w:numId'))
                if numId_elem is not None:
                    target_numId = numId_elem.get(qn('w:val'))
                    break

    if target_numId is None:
        return None, None

    # Find abstractNumId for this numId
    target_absNumId = None
    for num in numbering_xml.findall(qn('w:num')):
        if num.get(qn('w:numId')) == target_numId:
            absRef = num.find(qn('w:abstractNumId'))
            if absRef is not None:
                target_absNumId = absRef.get(qn('w:val'))
            break

    if target_absNumId is None:
        return None, None

    # Extract bullet chars from abstractNum definition
    lvl0_char = None
    lvl1_char = None
    for absNum in numbering_xml.findall(qn('w:abstractNum')):
        if absNum.get(qn('w:abstractNumId')) == target_absNumId:
            for lvl in absNum.findall(qn('w:lvl')):
                ilvl = lvl.get(qn('w:ilvl'))
                numFmt = lvl.find(qn('w:numFmt'))
                lvlText = lvl.find(qn('w:lvlText'))
                if numFmt is not None and numFmt.get(qn('w:val')) == 'bullet':
                    char = lvlText.get(qn('w:val')) if lvlText is not None else None
                    if ilvl == '0':
                        lvl0_char = char
                    elif ilvl == '1':
                        lvl1_char = char

    # Also check for lvlOverride in the num element
    for num in numbering_xml.findall(qn('w:num')):
        if num.get(qn('w:numId')) == target_numId:
            for ov in num.findall(qn('w:lvlOverride')):
                ilvl = ov.get(qn('w:ilvl'))
                lvl = ov.find(qn('w:lvl'))
                if lvl is not None:
                    numFmt = lvl.find(qn('w:numFmt'))
                    lvlText = lvl.find(qn('w:lvlText'))
                    if numFmt is not None and numFmt.get(qn('w:val')) == 'bullet':
                        char = lvlText.get(qn('w:val')) if lvlText is not None else None
                        if ilvl == '0':
                            lvl0_char = char
                        elif ilvl == '1':
                            lvl1_char = char

    return lvl0_char, lvl1_char


def get_list_structure(doc):
    """Extract list paragraph info: (ilvl, text) for each list paragraph."""
    from docx.oxml.ns import qn

    items = []
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                ilvl_elem = numPr.find(qn('w:ilvl'))
                if ilvl_elem is not None:
                    ilvl = int(ilvl_elem.get(qn('w:val')))
                    items.append((ilvl, para.text.strip()))
    return items


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract numbering info
    try:
        lvl0_char, lvl1_char = get_numbering_info(doc)
        print(f"INFO: Level 0 bullet char: {repr(lvl0_char)} (U+{ord(lvl0_char):04X})" if lvl0_char else "INFO: Level 0 bullet char: None")
        print(f"INFO: Level 1 bullet char: {repr(lvl1_char)} (U+{ord(lvl1_char):04X})" if lvl1_char else "INFO: Level 1 bullet char: None")
    except Exception as e:
        print(f"ERROR: Could not extract numbering info: {e}")
        lvl0_char, lvl1_char = None, None

    # Component 1: Level-1 bullets use square character (0.5 points)
    # This is THE task-introduced change: level-1 bullets must be square
    try:
        if lvl1_char is not None and lvl1_char in SQUARE_BULLETS:
            print(f"PASS: Component 1 — Level-1 bullet is square (U+{ord(lvl1_char):04X}) (0.5 pts)")
            total_score += 0.5
        else:
            actual = f"U+{ord(lvl1_char):04X}" if lvl1_char else "None"
            print(f"FAIL: Component 1 — Expected square bullet for level-1, found: {actual}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Level-0 bullets remain round AND level-1 is square (0.3 points)
    # Compound check: both conditions must be true to score
    # This ensures the task was done selectively (only level-1 changed)
    try:
        lvl0_is_round = lvl0_char is not None and lvl0_char in ROUND_BULLETS
        lvl1_is_square = lvl1_char is not None and lvl1_char in SQUARE_BULLETS
        if lvl0_is_round and lvl1_is_square:
            print(f"PASS: Component 2 — Level-0 round (U+{ord(lvl0_char):04X}) + Level-1 square (U+{ord(lvl1_char):04X}) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Level-0 round={lvl0_is_round}, Level-1 square={lvl1_is_square}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text content and hierarchy preserved AND level-1 is square (0.2 points)
    # Compound check anchored to the change: hierarchy must be intact AND bullet change applied
    try:
        items = get_list_structure(doc)
        lvl0_count = sum(1 for ilvl, _ in items if ilvl == 0)
        lvl1_count = sum(1 for ilvl, _ in items if ilvl == 1)
        hierarchy_ok = lvl0_count == 4 and lvl1_count == 7
        lvl1_square = lvl1_char is not None and lvl1_char in SQUARE_BULLETS

        if hierarchy_ok and lvl1_square:
            print(f"PASS: Component 3 — Hierarchy preserved (4 L0, 7 L1) + Level-1 square (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Hierarchy: L0={lvl0_count}/4, L1={lvl1_count}/7, square={lvl1_square}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
