"""
Initial Setup: Magazine-style article with single-column body section
Task ID: osworld_writer_section_columns_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_section_columns_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ---- Section 0: Header section (single column, title area) ----
    # Use the default first section for the title/header
    section0 = doc.sections[0]
    section0.page_width = Inches(8.5)
    section0.page_height = Inches(11)
    section0.left_margin = Inches(1.0)
    section0.right_margin = Inches(1.0)
    section0.top_margin = Inches(1.0)
    section0.bottom_margin = Inches(1.0)

    # Magazine title
    title = doc.add_heading('The Future of Urban Mobility', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)

    # Subtitle
    subtitle = doc.add_paragraph('How Smart Cities Are Reinventing the Way We Move')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Byline
    byline = doc.add_paragraph()
    byline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = byline.add_run('By Dr. Amelia Hartman, Urban Planning Institute')
    run.font.size = Pt(11)
    run.font.bold = True

    # Date line
    dateline = doc.add_paragraph()
    dateline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dateline.add_run('Published: March 2025  |  Special Feature')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Horizontal rule (simulated with underscores paragraph)
    divider = doc.add_paragraph('_' * 80)
    divider.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in divider.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ---- Section 1: Body section (single column — agent will change to 2 cols) ----
    # Add a continuous section break to start the body section
    body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    body_section.left_margin = Inches(1.0)
    body_section.right_margin = Inches(1.0)
    body_section.top_margin = Inches(0.5)
    body_section.bottom_margin = Inches(1.0)

    # Ensure body section is single column (default; do NOT set multi-column here)
    # The body_section.sectPr should NOT contain <w:cols> with num > 1

    # First body paragraph — intro
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'Across the globe, cities are grappling with unprecedented growth. '
        'From São Paulo to Seoul, urban planners face mounting pressure to '
        'deliver transportation networks that are efficient, sustainable, and '
        'inclusive. The old model — highways, parking lots, and dependence on '
        'private automobiles — is giving way to a bold new vision of integrated '
        'mobility hubs and on-demand transit.'
    )
    intro_run.font.size = Pt(11)
    intro.paragraph_format.space_after = Pt(8)

    # Section heading
    h2 = doc.add_heading('The Rise of Mobility-as-a-Service', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)

    p2 = doc.add_paragraph()
    p2.add_run(
        'Mobility-as-a-Service (MaaS) platforms are fundamentally changing the '
        'commuter experience. Rather than owning a car, city residents in Helsinki, '
        'Singapore, and Los Angeles can now plan, book, and pay for multimodal trips '
        'through a single app. Integrating bus, metro, ride-share, bicycle hire, '
        'and even electric scooters into one seamless journey reduces both cost and '
        'carbon emissions.'
    ).font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(8)

    p3 = doc.add_paragraph()
    p3.add_run(
        'Helsinki\'s Whim application, launched in 2017, became the world\'s first '
        'fully operational MaaS platform. Subscription tiers allow users to pay a '
        'flat monthly fee for unlimited transit access — including taxis for short '
        'distances after midnight. Usage data from 2024 shows a 12% decline in '
        'private car registrations among Whim subscribers.'
    ).font.size = Pt(11)
    p3.paragraph_format.space_after = Pt(8)

    # Section heading
    h3 = doc.add_heading('Autonomous Vehicles and Last-Mile Connectivity', level=2)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)

    p4 = doc.add_paragraph()
    p4.add_run(
        'Autonomous vehicle (AV) technology has advanced rapidly, with Level 4 '
        'self-driving shuttles now operating in geofenced zones across Phoenix, '
        'Shenzhen, and Dubai. These small electric pods — typically carrying six '
        'to twelve passengers — bridge the gap between metro stations and '
        'residential neighborhoods, solving the perennial "last mile" problem '
        'that has long hampered public transit adoption.'
    ).font.size = Pt(11)
    p4.paragraph_format.space_after = Pt(8)

    p5 = doc.add_paragraph()
    p5.add_run(
        'Safety records are improving year-on-year. The Waymo One service logged '
        'over 10 million driverless miles in 2024 with zero fatalities. Critics '
        'remain cautious, pointing to edge cases in adverse weather and the '
        'ethical complexities of algorithmic decision-making. Regulatory frameworks '
        'are still evolving in most jurisdictions, creating an uneven patchwork '
        'of deployment rules.'
    ).font.size = Pt(11)
    p5.paragraph_format.space_after = Pt(8)

    # Section heading
    h4 = doc.add_heading('Cycling Infrastructure Renaissance', level=2)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)

    p6 = doc.add_paragraph()
    p6.add_run(
        'Protected cycling lanes — physically separated from vehicle traffic — '
        'have exploded in popularity since the 2020 pandemic. Paris added over '
        '1,000 km of new cycling infrastructure between 2020 and 2024, and cycling '
        'modal share in the city center reached 22%. Amsterdam, long a global '
        'cycling champion, is now investing in underground bicycle parking '
        'at major rail stations to reduce surface congestion.'
    ).font.size = Pt(11)
    p6.paragraph_format.space_after = Pt(8)

    p7 = doc.add_paragraph()
    p7.add_run(
        'The economic case for cycling is compelling. A study by the European '
        'Cyclists\' Federation estimates that shifting 20% of urban trips to '
        'bicycle would reduce EU healthcare costs by €50 billion annually. '
        'E-bikes are accelerating the transition, extending the practical range '
        'of cycling to 15–25 km and making hills and headwinds a non-issue for '
        'everyday commuters.'
    ).font.size = Pt(11)
    p7.paragraph_format.space_after = Pt(8)

    # Section heading
    h5 = doc.add_heading('Smart Traffic Management', level=2)
    for run in h5.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)

    p8 = doc.add_paragraph()
    p8.add_run(
        'Artificial intelligence is transforming how cities manage traffic flow. '
        'Adaptive signal control systems — which use real-time sensor data and '
        'machine learning to optimize traffic light timing — are reducing urban '
        'congestion by up to 30% in pilot cities. Pittsburgh\'s SURTRAC system '
        'demonstrated a 25% reduction in travel time and a 21% drop in vehicle '
        'emissions in its first year of operation.'
    ).font.size = Pt(11)
    p8.paragraph_format.space_after = Pt(8)

    p9 = doc.add_paragraph()
    p9.add_run(
        'Data privacy remains a significant concern. Smart city sensors collect '
        'vast quantities of movement data, raising questions about surveillance '
        'and civil liberties. Cities like Barcelona have adopted open-data '
        'frameworks to ensure transparency, while others are exploring federated '
        'learning techniques to train AI models without centralizing personal '
        'location data.'
    ).font.size = Pt(11)
    p9.paragraph_format.space_after = Pt(8)

    # Closing paragraph
    conclusion = doc.add_paragraph()
    conclusion.add_run(
        'The transformation of urban mobility is well underway. What emerges '
        'over the next decade will depend on political will, investment in '
        'equitable infrastructure, and the ability of technologists and '
        'planners to collaborate across disciplines. One thing is clear: '
        'the city of the future will move very differently from the city of today.'
    ).font.size = Pt(11)
    conclusion.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
