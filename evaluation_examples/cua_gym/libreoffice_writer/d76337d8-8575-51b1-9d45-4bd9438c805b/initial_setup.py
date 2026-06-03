"""
Initial Setup: Regex Find & Replace to wrap standalone numbers with square brackets
Task ID: writer_edit_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'data_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Annual Data Report — FY 2024', level=0)

    # Section 1 — Overview
    doc.add_heading('Section 1: Overview', level=1)
    doc.add_paragraph(
        'This report summarises the key findings from 42 independent surveys '
        'conducted across 7 regional offices during the fiscal year. A total of '
        '3 executive sponsors reviewed the methodology before approval.'
    )
    doc.add_paragraph(
        'Response rates improved by 18 percentage points compared to the prior '
        'cycle, with 312 valid submissions received out of 350 distributed.'
    )

    # Section 2 — Participant Demographics
    doc.add_heading('Section 2: Participant Demographics', level=1)
    doc.add_paragraph(
        'There were 42 participants enrolled in the primary cohort. Of these, '
        '15 were from the engineering division, 12 from product management, '
        '9 from sales, and 6 from operations.'
    )
    doc.add_paragraph(
        'The age distribution ranged from 24 to 58 years, with a median age of '
        '36. Tenure across all participants averaged 5 years, with a range of '
        '1 to 22 years of service.'
    )
    # Paragraph with chemical formula — H2O and F16 should not be changed
    doc.add_paragraph(
        'Lab samples were tested using H2O as the baseline solvent. The F16 '
        'aircraft component stress test was performed at 40 kPa over a period '
        'of 72 hours. Compound B2O3 was excluded due to contamination in 3 '
        'of the 8 batches.'
    )

    # Section 3 — Financial Summary
    doc.add_heading('Section 3: Financial Summary', level=1)
    doc.add_paragraph(
        'Section 3 covers the full breakdown of departmental spending. '
        'The cost was 150 dollars per unit for the first 200 units, '
        'dropping to 120 dollars per unit for orders above 500 units.'
    )
    doc.add_paragraph(
        'Total expenditure reached 2450000 dollars, exceeding the initial '
        'budget of 2200000 by approximately 11 percent. An emergency reserve '
        'of 100000 dollars was authorised by the board on day 45 of the project.'
    )

    # Section 4 — Timeline
    doc.add_heading('Section 4: Timeline and Milestones', level=1)
    doc.add_paragraph(
        'Phase 1 was completed in 90 days, 5 days ahead of schedule. '
        'Phase 2 required an extension of 14 days due to supply chain delays '
        'affecting 3 critical vendors. The final delivery occurred on day 180.'
    )
    doc.add_paragraph(
        'A total of 28 milestones were tracked using the project management '
        'dashboard. Of these, 25 were completed on time and 3 required '
        'rescheduling. No milestones were cancelled.'
    )

    # Section 5 — Recommendations
    doc.add_heading('Section 5: Recommendations', level=1)
    doc.add_paragraph(
        'Based on the analysis, 6 recommendations have been identified for '
        'the upcoming fiscal year. Priority 1 focuses on increasing headcount '
        'by at least 10 full-time equivalents in the engineering team.'
    )
    doc.add_paragraph(
        'Recommendation 2 calls for a reduction in vendor contracts from 17 '
        'to no more than 9, consolidating spend and reducing administrative '
        'overhead by an estimated 30 percent.'
    )
    doc.add_paragraph(
        'Recommendation 3 proposes allocating 500000 dollars to upgrade '
        'the legacy infrastructure, which is currently running on hardware '
        'that is over 8 years old. This is considered critical for '
        'maintaining uptime above 99 percent.'
    )

    # Footer note
    doc.add_paragraph('')
    note = doc.add_paragraph(
        'Prepared by the Analytics Team | Document version 4 | Confidential'
    )
    note.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
