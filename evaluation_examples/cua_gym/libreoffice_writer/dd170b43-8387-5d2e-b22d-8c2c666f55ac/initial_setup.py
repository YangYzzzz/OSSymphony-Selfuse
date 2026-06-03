"""
Initial Setup: Mail merge template with data source for payment reminders
Task ID: writer_mt_022
Domain: libreoffice_writer
"""

import os
import csv
import shlex
import subprocess
import time
import random
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_022'
OUTPUT_DOCX = f'{WORKDIR}/{TASK_ID}.docx'
OUTPUT_CSV = f'{WORKDIR}/Invoices.csv'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_data_source():
    """Create CSV data source with 50 invoice records.
    Exactly 22 records have InvoiceAmount > 500.
    """
    random.seed(42)

    first_names = [
        "Sarah", "Marcus", "Elena", "David", "Priya",
        "James", "Mei", "Carlos", "Fatima", "Robert",
        "Aisha", "Thomas", "Yuki", "Patrick", "Nadia",
        "William", "Rosa", "Ahmed", "Laura", "Kevin",
        "Olivia", "Daniel", "Sonia", "Brian", "Anita",
        "Richard", "Yolanda", "Samuel", "Grace", "Victor",
        "Hannah", "Oscar", "Diana", "Frank", "Lena",
        "George", "Tanya", "Ivan", "Maria", "Nathan",
        "Chloe", "Peter", "Sylvia", "Alex", "Beatrice",
        "Henry", "Julia", "Martin", "Wendy", "Stephen"
    ]

    last_names = [
        "Chen", "Johnson", "Kowalski", "Patel", "Sharma",
        "Williams", "Tanaka", "Rodriguez", "Al-Rashid", "Mitchell",
        "Okafor", "Anderson", "Suzuki", "O'Brien", "Ivanova",
        "Thompson", "Gonzalez", "Hassan", "Fischer", "Park",
        "Bennett", "Torres", "Mehta", "Collins", "Gupta",
        "Morrison", "Cruz", "Adler", "Nakamura", "Sullivan",
        "Foster", "Vargas", "Petrov", "Reeves", "Larsen",
        "Hayes", "Volkov", "Kato", "Fernandez", "Douglas",
        "Webb", "Olsen", "Singh", "Maxwell", "Laurent",
        "Price", "Hoffman", "Novak", "Yamamoto", "Barnes"
    ]

    # Generate exactly 22 amounts > 500 and 28 amounts <= 500
    high_amounts = []
    while len(high_amounts) < 22:
        amt = round(random.uniform(501, 5000), 2)
        high_amounts.append(amt)

    low_amounts = []
    while len(low_amounts) < 28:
        amt = round(random.uniform(75, 500), 2)
        low_amounts.append(amt)

    all_amounts = high_amounts + low_amounts
    random.shuffle(all_amounts)

    # Generate invoice numbers and due dates
    months = ['01', '02', '03', '04', '05', '06', '07', '08', '09', '10', '11', '12']
    days_range = range(1, 29)
    base_inv = 10000

    records = []
    for i in range(50):
        client_name = f"{first_names[i]} {last_names[i]}"
        invoice_number = f"INV-{base_inv + i + 1}"
        invoice_amount = all_amounts[i]
        month = months[i % 12]
        day = random.choice(list(days_range))
        due_date = f"2025-{month}-{day:02d}"
        records.append([client_name, invoice_number, invoice_amount, due_date])

    with open(OUTPUT_CSV, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['ClientName', 'InvoiceNumber', 'InvoiceAmount', 'DueDate'])
        for record in records:
            writer.writerow(record)

    print(f"Data source created: {OUTPUT_CSV} ({len(records)} records)")

    # Count verification
    count_gt500 = sum(1 for r in records if r[2] > 500)
    print(f"Records with InvoiceAmount > 500: {count_gt500}")
    return records


def create_template_document():
    """Create the mail merge template document with merge field placeholders."""
    doc = Document()

    # Document title
    title = doc.add_heading('Payment Reminder', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = date_para.add_run('Date: March 15, 2025')
    run.font.size = Pt(11)

    # Blank line
    doc.add_paragraph()

    # Recipient
    dear_para = doc.add_paragraph()
    run = dear_para.add_run('Dear ')
    run.font.size = Pt(11)
    run = dear_para.add_run('<ClientName>')
    run.font.size = Pt(11)
    run.bold = True
    run = dear_para.add_run(',')
    run.font.size = Pt(11)

    # Body paragraph 1
    body1 = doc.add_paragraph()
    run = body1.add_run(
        'This letter serves as a formal reminder that your invoice is approaching its due date. '
        'Please find the details of your outstanding payment below:'
    )
    run.font.size = Pt(11)

    # Invoice details
    doc.add_paragraph()
    details = [
        ('Invoice Number: ', '<InvoiceNumber>'),
        ('Invoice Amount: $', '<InvoiceAmount>'),
        ('Due Date: ', '<DueDate>'),
    ]
    for label, field in details:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(label)
        run.font.size = Pt(11)
        run = p.add_run(field)
        run.font.size = Pt(11)
        run.bold = True

    # Body paragraph 2
    doc.add_paragraph()
    body2 = doc.add_paragraph()
    run = body2.add_run(
        'We kindly request that you process this payment by the due date indicated above. '
        'If you have already submitted your payment, please disregard this notice. '
        'Should you have any questions regarding this invoice, please do not hesitate to '
        'contact our Accounts Receivable department at accounts@meridianservices.com or '
        'call us at (555) 234-8901.'
    )
    run.font.size = Pt(11)

    # Closing
    doc.add_paragraph()
    closing = doc.add_paragraph()
    run = closing.add_run('Sincerely,')
    run.font.size = Pt(11)

    doc.add_paragraph()

    sig = doc.add_paragraph()
    run = sig.add_run('Rachel Thompson')
    run.font.size = Pt(11)
    run.bold = True

    sig_title = doc.add_paragraph()
    run = sig_title.add_run('Accounts Receivable Manager')
    run.font.size = Pt(11)

    company = doc.add_paragraph()
    run = company.add_run('Meridian Business Services, LLC')
    run.font.size = Pt(11)

    doc.save(OUTPUT_DOCX)
    print(f"Template document created: {OUTPUT_DOCX}")


def main():
    create_data_source()
    create_template_document()

    # Open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT_DOCX}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
