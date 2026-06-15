"""
Initial Setup: Insert date and author fields in document header
Task ID: writer_struct_056
Domain: libreoffice_writer

Creates a 6-page laboratory report (lab_notebook.docx) with:
- Empty but enabled header
- Document author property set to 'Dr. Elena Rivera'
- Realistic lab notebook body content
- Opens in LibreOffice Writer for GUI agent
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'lab_notebook'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_core_properties(doc, author='Dr. Elena Rivera', title='Laboratory Research Notebook'):
    """Set document core properties including author."""
    core_props = doc.core_properties
    core_props.author = author
    core_props.title = title
    core_props.subject = 'Biochemistry Research'
    core_props.keywords = 'laboratory, research, biochemistry, experiments'


def add_heading(doc, text, level=1):
    """Add a heading paragraph."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def create_initial():
    doc = Document()

    # Set document core properties (author = Dr. Elena Rivera)
    set_core_properties(doc, author='Dr. Elena Rivera', title='Laboratory Research Notebook')

    # Configure section - enable header but leave it empty
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1.0)

    # Enable header (but leave empty - task requires agent to add content)
    header = section.header
    header.is_linked_to_previous = False
    # Leave header paragraph empty (it exists by default as empty paragraph)
    # Ensure the header paragraph is blank
    if header.paragraphs:
        for para in header.paragraphs:
            for run in para.runs:
                run.text = ''
            # Remove all runs, keep empty paragraph
            for run in para.runs:
                run._element.getparent().remove(run._element)

    # ----------------------------------------------------------------
    # PAGE 1: Title page
    # ----------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(48)
    title_run = title_para.add_run('Laboratory Research Notebook')
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('Biochemistry & Molecular Biology Division')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info_para.add_run('Principal Investigator: Dr. Elena Rivera\n')
    info_para.add_run('Department of Biochemistry, Westfield University\n')
    info_para.add_run('Project Code: WU-BIO-2025-041\n')
    info_para.add_run('Notebook Period: January – June 2025')

    doc.add_page_break()

    # ----------------------------------------------------------------
    # PAGE 2: Table of Contents / Overview
    # ----------------------------------------------------------------
    add_heading(doc, 'Overview and Objectives', level=1)

    add_paragraph(doc,
        'This laboratory notebook documents ongoing research into the enzymatic degradation of '
        'polyethylene terephthalate (PET) plastics using engineered variants of PETase derived '
        'from Ideonella sakaiensis. The primary objective is to optimize reaction conditions '
        'for maximum substrate turnover at ambient temperatures.')

    add_paragraph(doc,
        'Secondary objectives include characterization of novel variant binding kinetics, '
        'structural analysis via cryo-EM, and assessment of co-substrate requirements. '
        'All experiments adhere to institutional biosafety protocols (BSL-1).')

    add_heading(doc, 'Project Team', level=2)

    # Team table
    team_table = doc.add_table(rows=5, cols=3)
    team_table.style = 'Table Grid'
    headers = ['Name', 'Role', 'Contact']
    for col_idx, header_text in enumerate(headers):
        cell = team_table.cell(0, col_idx)
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True

    team_data = [
        ['Dr. Elena Rivera', 'Principal Investigator', 'erivera@westfield.edu'],
        ['Marcus Chen', 'Research Associate', 'mchen@westfield.edu'],
        ['Priya Nair', 'PhD Candidate', 'pnair@westfield.edu'],
        ['James Okonkwo', 'Lab Technician', 'jokonkwo@westfield.edu'],
    ]
    for row_idx, row_data in enumerate(team_data, 1):
        for col_idx, cell_text in enumerate(row_data):
            team_table.cell(row_idx, col_idx).text = cell_text

    doc.add_page_break()

    # ----------------------------------------------------------------
    # PAGE 3: Experiment Log - January
    # ----------------------------------------------------------------
    add_heading(doc, 'Experiment Log — January 2025', level=1)

    add_heading(doc, 'Experiment 001: Baseline PETase Activity Assay', level=2)

    add_paragraph(doc, 'Date: January 8, 2025', bold=True)
    add_paragraph(doc, 'Operator: Marcus Chen')
    add_paragraph(doc, 'Protocol Reference: WU-BIO-PROT-014 v2.1')

    add_heading(doc, 'Materials', level=3)
    doc.add_paragraph('PET film substrate (Goodfellow, 250 μm thickness)', style='List Bullet')
    doc.add_paragraph('Wild-type PETase (purified in-house, batch #WT-2024-08)', style='List Bullet')
    doc.add_paragraph('Glycine-NaOH buffer (pH 9.0, 50 mM)', style='List Bullet')
    doc.add_paragraph('HPLC-grade methanol and acetonitrile', style='List Bullet')
    doc.add_paragraph('Terephthalic acid (TPA) standard (Sigma-Aldrich, ≥99%)', style='List Bullet')

    add_heading(doc, 'Procedure', level=3)
    add_paragraph(doc,
        'PET film was cut into 1 cm² pieces and pre-washed with 70% ethanol followed by '
        'deionized water. Each piece was incubated in a 1 mL reaction volume containing '
        '100 nM enzyme in glycine-NaOH buffer at 30 °C with orbital shaking at 150 rpm. '
        'Reactions were terminated at 24, 48, and 72 hours by transferring to 95 °C for '
        '10 minutes. TPA release was quantified by HPLC.')

    add_heading(doc, 'Results', level=3)

    # Results table
    results_table = doc.add_table(rows=4, cols=4)
    results_table.style = 'Table Grid'
    res_headers = ['Time (h)', 'TPA Released (μM)', 'BHET (μM)', 'Conversion (%)']
    for col_idx, h in enumerate(res_headers):
        run = results_table.cell(0, col_idx).paragraphs[0].add_run(h)
        run.bold = True

    results_data = [
        ['24', '12.3 ± 0.8', '3.1 ± 0.2', '2.1'],
        ['48', '28.7 ± 1.4', '5.8 ± 0.4', '4.9'],
        ['72', '51.2 ± 2.1', '9.2 ± 0.6', '8.7'],
    ]
    for row_idx, row_data in enumerate(results_data, 1):
        for col_idx, val in enumerate(row_data):
            results_table.cell(row_idx, col_idx).text = val

    doc.add_page_break()

    # ----------------------------------------------------------------
    # PAGE 4: Experiment Log - February
    # ----------------------------------------------------------------
    add_heading(doc, 'Experiment Log — February 2025', level=1)

    add_heading(doc, 'Experiment 007: Variant IsPETase-S290P/C264A Characterization', level=2)

    add_paragraph(doc, 'Date: February 3, 2025', bold=True)
    add_paragraph(doc, 'Operator: Priya Nair')

    add_paragraph(doc,
        'Thermostability measurements were conducted using differential scanning fluorimetry (DSF). '
        'The double mutant S290P/C264A showed a significant improvement in thermal stability '
        'compared to wild-type, with a melting temperature (Tm) of 62.4 °C versus 46.1 °C for WT.')

    add_heading(doc, 'Experiment 008: pH Optimization Screen', level=2)
    add_paragraph(doc, 'Date: February 17, 2025', bold=True)
    add_paragraph(doc, 'Operator: Marcus Chen')

    add_paragraph(doc,
        'A pH range screen from 6.0 to 10.0 was conducted using Britton-Robinson buffer system. '
        'Optimal activity was observed at pH 9.0 consistent with literature values. Activity '
        'dropped sharply below pH 7.5 and above pH 9.5, indicating narrow pH tolerance.')

    # pH results table
    ph_table = doc.add_table(rows=6, cols=3)
    ph_table.style = 'Table Grid'
    ph_headers = ['pH', 'Relative Activity (%)', 'Notes']
    for col_idx, h in enumerate(ph_headers):
        run = ph_table.cell(0, col_idx).paragraphs[0].add_run(h)
        run.bold = True

    ph_data = [
        ['6.0', '8.2', 'Low activity'],
        ['7.0', '34.7', 'Moderate'],
        ['8.0', '76.3', 'Good'],
        ['9.0', '100.0', 'Optimal'],
        ['10.0', '52.1', 'Reduced'],
    ]
    for row_idx, row_data in enumerate(ph_data, 1):
        for col_idx, val in enumerate(row_data):
            ph_table.cell(row_idx, col_idx).text = val

    doc.add_page_break()

    # ----------------------------------------------------------------
    # PAGE 5: Structural Analysis
    # ----------------------------------------------------------------
    add_heading(doc, 'Structural Analysis', level=1)

    add_heading(doc, 'Cryo-EM Sample Preparation', level=2)
    add_paragraph(doc, 'Date: March 5, 2025', bold=True)
    add_paragraph(doc, 'Operator: Dr. Elena Rivera')

    add_paragraph(doc,
        'Purified IsPETase-S290P/C264A at 2 mg/mL was applied to glow-discharged Quantifoil '
        'R1.2/1.3 holey carbon grids. Grids were plunge-frozen in liquid ethane using a '
        'Vitrobot Mark IV (blot force 3, blot time 4 s, 100% humidity, 4 °C). '
        'Data collection was performed at the National Cryo-EM Facility on a Titan Krios '
        'equipped with a Falcon 4i detector at 300 kV.')

    add_heading(doc, 'Data Collection Parameters', level=3)

    cryo_table = doc.add_table(rows=7, cols=2)
    cryo_table.style = 'Table Grid'
    cryo_headers = ['Parameter', 'Value']
    for col_idx, h in enumerate(cryo_headers):
        run = cryo_table.cell(0, col_idx).paragraphs[0].add_run(h)
        run.bold = True

    cryo_data = [
        ['Microscope', 'Titan Krios G3i'],
        ['Voltage', '300 kV'],
        ['Detector', 'Falcon 4i (counting mode)'],
        ['Pixel size', '0.824 Å/pixel'],
        ['Total dose', '50 e⁻/Å²'],
        ['Defocus range', '−1.0 to −3.0 μm'],
    ]
    for row_idx, row_data in enumerate(cryo_data, 1):
        for col_idx, val in enumerate(row_data):
            cryo_table.cell(row_idx, col_idx).text = val

    add_heading(doc, 'Preliminary Structural Results', level=3)
    add_paragraph(doc,
        'Initial 2D class averages showed clear secondary structure features consistent with '
        'a globular enzyme with visible alpha-helices. 3D reconstruction at 2.8 Å resolution '
        'revealed a modified active site geometry in the double mutant with altered loop '
        'conformation around residues 285–295, potentially explaining the increased thermal stability.')

    doc.add_page_break()

    # ----------------------------------------------------------------
    # PAGE 6: Conclusions and Next Steps
    # ----------------------------------------------------------------
    add_heading(doc, 'Conclusions and Next Steps', level=1)

    add_heading(doc, 'Summary of Key Findings (Jan–Mar 2025)', level=2)

    doc.add_paragraph(
        'Wild-type PETase achieves 8.7% PET conversion in 72 h under baseline conditions.',
        style='List Number')
    doc.add_paragraph(
        'The S290P/C264A double mutant exhibits a 16.3 °C increase in melting temperature.',
        style='List Number')
    doc.add_paragraph(
        'Optimal reaction pH is 9.0; activity is sensitive to pH deviations > 1 unit.',
        style='List Number')
    doc.add_paragraph(
        'Cryo-EM structure at 2.8 Å resolution reveals modified active site loop conformation.',
        style='List Number')

    add_heading(doc, 'Planned Experiments (April–June 2025)', level=2)

    add_paragraph(doc,
        'Future work will focus on engineering additional stabilizing mutations identified '
        'from the cryo-EM structure, including potential disulfide bonds and hydrophobic '
        'core packing improvements. Kinetic characterization using stopped-flow fluorescence '
        'will provide kcat and KM values for wild-type and variants.')

    add_paragraph(doc,
        'Collaborative project with the Materials Science department will explore surface '
        'functionalization of PET films to enhance enzyme accessibility. Industrial '
        'scale-up feasibility assessment is planned for Q3 2025.')

    add_heading(doc, 'References', level=2)

    refs = [
        'Yoshida, S. et al. (2016) A bacterium that degrades and assimilates poly(ethylene terephthalate). Science 351, 1196–1199.',
        'Austin, H.P. et al. (2018) Characterization and engineering of a plastic-degrading aromatic polyesterase. PNAS 115, E4350–E4357.',
        'Tournier, V. et al. (2020) An engineered PET depolymerase to break down and recycle plastic bottles. Nature 580, 216–219.',
        'Lu, H. et al. (2022) Machine learning-aided engineering of hydrolases for PET depolymerization. Nature 604, 662–667.',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Number')

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
