"""
Initial Setup: Add page numbers to footer of a document (pre-task state)
Task ID: osworld_writer_easy_017
Domain: libreoffice_writer

Creates a 6-page chapter draft document with NO footer/page numbers.
The task is to add centered page numbers to the footer.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_easy_017'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set A4 page size
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Ensure footer is empty (no page numbers) — this is the initial state
    footer = section.footer
    footer.is_linked_to_previous = False
    # Leave footer empty (no content) — task requires adding page numbers

    # -----------------------------------------------------------------------
    # Chapter 1: Introduction to Digital Marketing
    # -----------------------------------------------------------------------
    heading1 = doc.add_heading('Chapter 1: Introduction to Digital Marketing', level=1)

    doc.add_paragraph(
        'Digital marketing has fundamentally transformed how businesses connect with their '
        'customers. Unlike traditional marketing methods that rely on physical media such as '
        'print advertisements and billboards, digital marketing leverages the power of the '
        'internet and electronic devices to reach a wider, more targeted audience.'
    )

    doc.add_paragraph(
        'The evolution of digital marketing can be traced back to the early days of the '
        'internet in the 1990s. As more households gained internet access, businesses began '
        'recognizing the potential of online advertising. Early strategies included banner ads '
        'and basic email campaigns, which laid the foundation for the sophisticated digital '
        'marketing ecosystems we see today.'
    )

    doc.add_heading('1.1 Key Digital Marketing Channels', level=2)

    doc.add_paragraph(
        'Modern digital marketing encompasses a diverse array of channels, each offering '
        'unique advantages and challenges. Search Engine Optimization (SEO) remains one of '
        'the most cost-effective long-term strategies, helping businesses achieve organic '
        'visibility on platforms like Google and Bing.'
    )

    doc.add_paragraph(
        'Social media marketing has emerged as a dominant force, with platforms such as '
        'Instagram, Facebook, LinkedIn, and TikTok offering unprecedented access to '
        'highly segmented audience groups. Brands like Patagonia, Nike, and Airbnb have '
        'built strong digital communities through consistent, authentic social engagement.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Chapter 2: Content Strategy and Creation
    # -----------------------------------------------------------------------
    doc.add_heading('Chapter 2: Content Strategy and Creation', level=1)

    doc.add_paragraph(
        'Content is often described as the cornerstone of any successful digital marketing '
        'campaign. A well-defined content strategy ensures that all published material '
        'aligns with business objectives, resonates with the target audience, and delivers '
        'measurable value.'
    )

    doc.add_paragraph(
        'Before creating content, marketers must invest time in understanding their audience '
        'through detailed persona development. This involves researching demographic data, '
        'psychographic profiles, behavioral patterns, and pain points. Tools like Google '
        'Analytics, HubSpot, and Semrush provide valuable data insights that inform '
        'persona creation.'
    )

    doc.add_heading('2.1 The Content Creation Process', level=2)

    doc.add_paragraph(
        'Effective content creation follows a structured workflow that begins with ideation '
        'and keyword research. Content teams at organizations like HubSpot and Moz use '
        'keyword tools to identify topics that balance search volume with competitive '
        'difficulty, ensuring content can rank and drive organic traffic.'
    )

    doc.add_paragraph(
        'The writing phase demands clarity, accuracy, and engagement. Long-form blog posts '
        'between 1,500 and 3,000 words tend to perform well in search rankings, while '
        'shorter pieces of 500–800 words are ideal for social media syndication. Video '
        'content, particularly short-form videos under 60 seconds, has seen explosive '
        'growth across platforms like YouTube Shorts and Instagram Reels.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Chapter 3: Search Engine Optimization
    # -----------------------------------------------------------------------
    doc.add_heading('Chapter 3: Search Engine Optimization', level=1)

    doc.add_paragraph(
        'Search Engine Optimization (SEO) is the practice of optimizing web content to '
        'improve its visibility in organic search engine results. Given that approximately '
        '93% of online experiences begin with a search engine, mastering SEO is critical '
        'for any organization seeking sustainable online growth.'
    )

    doc.add_paragraph(
        'Google\'s algorithm updates — including Panda, Penguin, Hummingbird, and the '
        'more recent BERT and MUM models — have progressively shifted SEO from keyword '
        'stuffing tactics toward genuine user value creation. Modern SEO demands technical '
        'excellence, authoritative content, and strong link-building strategies.'
    )

    doc.add_heading('3.1 On-Page Optimization Techniques', level=2)

    doc.add_paragraph(
        'On-page SEO refers to all optimizations made directly within the webpage. This '
        'includes crafting compelling title tags (50–60 characters), writing descriptive '
        'meta descriptions (150–160 characters), using header hierarchy (H1 through H6) '
        'to structure content, and incorporating LSI keywords naturally throughout the text.'
    )

    doc.add_paragraph(
        'Page load speed is a critical ranking factor measured by Google\'s Core Web Vitals. '
        'Achieving a Largest Contentful Paint (LCP) under 2.5 seconds, a First Input Delay '
        '(FID) under 100 milliseconds, and a Cumulative Layout Shift (CLS) score below 0.1 '
        'are the benchmarks set by Google for excellent user experience.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Chapter 4: Social Media Marketing
    # -----------------------------------------------------------------------
    doc.add_heading('Chapter 4: Social Media Marketing', level=1)

    doc.add_paragraph(
        'Social media marketing involves creating and sharing content on social media '
        'networks to achieve marketing and branding goals. With over 4.9 billion social '
        'media users worldwide as of 2024, the potential reach of social media campaigns '
        'is unparalleled in the history of advertising.'
    )

    doc.add_paragraph(
        'Each social platform has its own unique culture, user demographics, and content '
        'formats. LinkedIn caters predominantly to B2B professionals, with average user '
        'ages between 25–34 years and a strong focus on thought leadership content. '
        'Instagram attracts younger audiences aged 18–29 who respond well to visually '
        'rich content, stories, and influencer collaborations.'
    )

    doc.add_heading('4.1 Building a Social Media Strategy', level=2)

    doc.add_paragraph(
        'A robust social media strategy begins with clearly defined goals. Whether the '
        'objective is brand awareness, lead generation, community building, or customer '
        'service, each goal requires distinct content approaches, posting cadences, and '
        'success metrics. SMART goal-setting frameworks — Specific, Measurable, Achievable, '
        'Relevant, and Time-bound — provide structure to strategy development.'
    )

    doc.add_paragraph(
        'Content calendars are essential tools for maintaining consistency and planning '
        'seasonal campaigns. Organizations like Buffer and Hootsuite offer social media '
        'management platforms that allow teams to schedule posts, monitor engagement '
        'metrics, and collaborate on content approval workflows across multiple accounts.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Chapter 5: Email Marketing
    # -----------------------------------------------------------------------
    doc.add_heading('Chapter 5: Email Marketing', level=1)

    doc.add_paragraph(
        'Despite the proliferation of new marketing channels, email marketing continues to '
        'deliver the highest return on investment of any digital channel, with industry '
        'studies consistently reporting returns of $36–$42 for every dollar spent. The '
        'direct, permission-based nature of email communication creates a uniquely personal '
        'connection between brands and their audiences.'
    )

    doc.add_paragraph(
        'Email list quality matters far more than list size. A curated list of 10,000 '
        'engaged subscribers will consistently outperform a list of 100,000 disengaged '
        'contacts. Regular list hygiene practices — removing inactive subscribers, '
        'correcting invalid addresses, and re-engaging dormant contacts — are essential '
        'for maintaining deliverability rates above 95%.'
    )

    doc.add_heading('5.1 Campaign Types and Automation', level=2)

    doc.add_paragraph(
        'Modern email marketing extends far beyond promotional broadcasts. Welcome series '
        'nurture new subscribers through onboarding workflows, introducing them to products '
        'and brand values over a sequence of carefully crafted emails. E-commerce companies '
        'leverage abandoned cart automations — triggered when a shopper adds items to their '
        'cart but does not complete the purchase — to recover an estimated 5–15% of lost revenue.'
    )

    doc.add_paragraph(
        'Behavioral segmentation enables marketers to deliver highly personalized content '
        'based on past purchase history, browsing behavior, geographic location, and '
        'lifecycle stage. Platforms like Klaviyo, Mailchimp, and ActiveCampaign provide '
        'sophisticated segmentation and A/B testing capabilities that allow continuous '
        'optimization of email performance metrics.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Chapter 6: Analytics and Performance Measurement
    # -----------------------------------------------------------------------
    doc.add_heading('Chapter 6: Analytics and Performance Measurement', level=1)

    doc.add_paragraph(
        'Data-driven decision making is the hallmark of effective digital marketing. '
        'Analytics platforms such as Google Analytics 4, Adobe Analytics, and Mixpanel '
        'provide comprehensive insights into user behavior, traffic sources, conversion '
        'funnels, and revenue attribution. Understanding how to interpret and act on '
        'these insights separates successful marketers from those who struggle to '
        'demonstrate ROI.'
    )

    doc.add_paragraph(
        'Key performance indicators (KPIs) vary by channel and objective. For SEO, '
        'organic traffic, keyword rankings, and domain authority are primary metrics. '
        'Paid advertising campaigns are evaluated on click-through rates, cost per click, '
        'conversion rates, and return on ad spend (ROAS). Email marketing performance '
        'centers on open rates, click rates, conversion rates, and unsubscribe rates.'
    )

    doc.add_heading('6.1 Attribution Models and Reporting', level=2)

    doc.add_paragraph(
        'Marketing attribution — determining which touchpoints deserve credit for driving '
        'conversions — remains one of the most complex challenges in digital marketing. '
        'Single-touch models like first-click and last-click attribution offer simplicity '
        'but fail to capture the multi-channel customer journeys typical of modern '
        'purchasing behavior.'
    )

    doc.add_paragraph(
        'Multi-touch attribution models such as linear, time-decay, and position-based '
        'distribute conversion credit across multiple touchpoints, providing a more '
        'accurate picture of the customer journey. Data-driven attribution, available in '
        'Google Analytics 4, uses machine learning algorithms to assign fractional credit '
        'based on the actual contribution of each touchpoint to conversions.'
    )

    doc.add_paragraph(
        'Building effective dashboards and reports requires translating raw data into '
        'actionable insights for stakeholders at various levels of the organization. '
        'Executive dashboards focus on high-level revenue and growth metrics, while '
        'channel-specific reports provide granular details for marketing teams optimizing '
        'day-to-day campaign performance. Tools like Looker Studio (formerly Google Data '
        'Studio), Tableau, and Power BI enable the creation of automated, visually '
        'compelling reports that keep teams aligned and informed.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
