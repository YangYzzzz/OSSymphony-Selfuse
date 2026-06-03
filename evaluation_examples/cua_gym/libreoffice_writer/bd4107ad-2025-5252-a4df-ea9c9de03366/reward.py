"""
Reward Script: Add page numbers to the footer (centered)
Task ID: osworld_writer_easy_017
Domain: libreoffice_writer
Scoring:
  Component 1: Footer contains a PAGE field code             (0.5 pts)
  Component 2: Footer paragraph is center-aligned            (0.3 pts)
  Component 3: PAGE field is properly formed (begin+instr+end)(0.2 pts)
  Total: 1.0
"""

import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_easy_017'


def has_page_field(para):
    """
    Check if a paragraph contains a PAGE field code.
    PAGE fields are encoded as:
      <w:fldChar fldCharType="begin"/>  <w:instrText> PAGE </w:instrText>  <w:fldChar fldCharType="end"/>
    Returns tuple: (has_page_instr, has_begin, has_end)
    """
    xml = para._element.xml
    has_page_instr = 'PAGE' in xml and 'instrText' in xml
    has_begin = 'fldCharType="begin"' in xml or "fldCharType='begin'" in xml
    has_end = 'fldCharType="end"' in xml or "fldCharType='end'" in xml
    return has_page_instr, has_begin, has_end


def verify_task(file_path):
    """
    Verify that the document footer contains centered page numbers.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — precondition gate
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section
    if not doc.sections:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    footer = section.footer

    # Get the first footer paragraph (the one that should hold the page number)
    if not footer.paragraphs:
        print("FAIL: Footer has no paragraphs at all")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    ftr_para = footer.paragraphs[0]

    # Component 1: Footer contains a PAGE field code (0.5 points)
    # This FAILS on initial (empty footer) and PASSES on golden (PAGE field present)
    try:
        has_page_instr, has_begin, has_end = has_page_field(ftr_para)
        if has_page_instr:
            print(f"PASS: Component 1 — PAGE field code found in footer (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — No PAGE instrText found in footer paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Footer paragraph is center-aligned (0.3 points)
    # This FAILS on initial (alignment=None) and PASSES on golden (CENTER)
    try:
        alignment = ftr_para.paragraph_format.alignment
        # Accept either CENTER enum or a jc/@val="center" in XML (both mean centered)
        is_centered = (alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
        # Also check XML directly for <w:jc w:val="center"/>
        if not is_centered:
            ftr_xml = ftr_para._element.xml
            is_centered = ('jc' in ftr_xml and 'center' in ftr_xml)
        if is_centered:
            print(f"PASS: Component 2 — Footer paragraph is center-aligned (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected CENTER alignment, got: {alignment}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: PAGE field is properly formed (begin + instrText + end) (0.2 points)
    # This FAILS on initial (no runs at all) and PASSES on golden (all 3 parts present)
    try:
        has_page_instr, has_begin, has_end = has_page_field(ftr_para)
        if has_page_instr and has_begin and has_end:
            print(f"PASS: Component 3 — PAGE field properly formed (begin + PAGE + end) (0.2 pts)")
            total_score += 0.2
        else:
            missing = []
            if not has_page_instr:
                missing.append("PAGE instrText")
            if not has_begin:
                missing.append("fldChar begin")
            if not has_end:
                missing.append("fldChar end")
            print(f"FAIL: Component 3 — Incomplete PAGE field, missing: {missing}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path in the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
