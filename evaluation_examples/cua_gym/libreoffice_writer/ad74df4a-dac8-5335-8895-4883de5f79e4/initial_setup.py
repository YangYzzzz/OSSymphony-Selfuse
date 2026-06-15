"""
Initial Setup: Insert a table of figures after the table of contents
Task ID: writer_tech_051
Domain: libreoffice_writer

Creates a technical document with a Table of Contents and 6 captioned figures
throughout. No Table of Figures exists yet.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_051'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc(doc):
    """Add a Table of Contents field to the document."""
    para = doc.add_paragraph()
    run = para.add_run()
    r_elem = run._element

    # Begin field
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r_elem.append(fldChar_begin)

    run2 = para.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instrText)

    run3 = para.add_run()
    fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._element.append(fldChar_separate)

    # Static placeholder text for TOC
    run4 = para.add_run("(Table of Contents - Update to populate)")
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run4.font.size = Pt(10)

    run5 = para.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._element.append(fldChar_end)


def add_figure_caption(doc, figure_num, caption_text):
    """Add an image placeholder and a proper SEQ-based figure caption."""
    # Add a simple placeholder rectangle as "figure"
    para_img = doc.add_paragraph()
    para_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_img = para_img.add_run("[Figure Placeholder]")
    run_img.font.size = Pt(11)
    run_img.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_img.italic = True

    # Caption paragraph using the "Caption" style approach
    # We create a paragraph that looks like: "Figure X: caption_text"
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_para.paragraph_format.space_after = Pt(12)

    # Apply the Caption style if available, otherwise format manually
    try:
        caption_para.style = doc.styles['Caption']
    except KeyError:
        pass

    # "Figure " label
    run_label = caption_para.add_run("Figure ")
    run_label.font.size = Pt(10)
    run_label.italic = True

    # SEQ field for auto-numbering: { SEQ Figure \* ARABIC }
    r_seq = caption_para.add_run()
    r_elem = r_seq._element

    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    r_elem.append(fldChar_begin)

    r_instr = caption_para.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> SEQ Figure \\* ARABIC </w:instrText>')
    r_instr._element.append(instrText)

    r_sep = caption_para.add_run()
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    r_sep._element.append(fldChar_sep)

    # Static number (will be updated by LO)
    r_num = caption_para.add_run(str(figure_num))
    r_num.font.size = Pt(10)
    r_num.italic = True

    r_end = caption_para.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    r_end._element.append(fldChar_end)

    # ": caption text"
    run_caption = caption_para.add_run(f": {caption_text}")
    run_caption.font.size = Pt(10)
    run_caption.italic = True

    return caption_para


def create_initial():
    doc = Document()

    # ---- Title Page ----
    title = doc.add_heading("Cloud Infrastructure Migration Guide", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Technical Reference Document\nVersion 2.4 - March 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("Prepared by: Infrastructure Engineering Team\nNexaCloud Solutions Inc.")
    run.font.size = Pt(11)

    doc.add_page_break()

    # ---- Table of Contents ----
    toc_heading = doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ---- Section 1: Introduction ----
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document provides a comprehensive guide for migrating on-premises "
        "infrastructure to cloud-based environments. The migration strategy outlined "
        "here covers assessment, planning, execution, and post-migration optimization "
        "phases. Organizations considering cloud migration should use this document "
        "as their primary technical reference throughout the transition process."
    )
    doc.add_paragraph(
        "The migration framework has been developed based on over 200 enterprise "
        "migrations completed between 2023 and 2025, incorporating lessons learned "
        "and best practices from diverse industry verticals including healthcare, "
        "financial services, and manufacturing."
    )

    doc.add_heading("1.1 Scope and Objectives", level=2)
    doc.add_paragraph(
        "The primary objectives of this migration initiative include reducing "
        "infrastructure costs by 35-40%, improving system availability to 99.95%, "
        "and enabling auto-scaling capabilities for peak demand periods. Secondary "
        "objectives include compliance with SOC 2 Type II and GDPR requirements."
    )

    # Figure 1
    add_figure_caption(doc, 1, "Migration timeline and key milestones for Phase 1 through Phase 4")

    # ---- Section 2: Current Architecture Assessment ----
    doc.add_heading("2. Current Architecture Assessment", level=1)
    doc.add_paragraph(
        "Before initiating migration, a thorough assessment of the existing "
        "infrastructure is essential. The assessment covers compute resources, "
        "storage systems, networking topology, and application dependencies. "
        "The NexaCloud Assessment Tool (NCAT) was deployed across all 47 production "
        "servers to collect telemetry data over a 30-day observation window."
    )

    doc.add_heading("2.1 Compute Resource Inventory", level=2)

    # Table: Server inventory
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Server Name", "CPU Cores", "RAM (GB)", "Avg. Utilization"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    server_data = [
        ["PROD-WEB-01", "16", "64", "72%"],
        ["PROD-WEB-02", "16", "64", "68%"],
        ["PROD-DB-01", "32", "256", "85%"],
        ["PROD-APP-01", "24", "128", "61%"],
        ["PROD-CACHE-01", "8", "32", "45%"],
    ]
    for r, row_data in enumerate(server_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph("")  # spacing

    # Figure 2
    add_figure_caption(doc, 2, "Current on-premises network topology with VLAN segmentation")

    doc.add_heading("2.2 Storage Analysis", level=2)
    doc.add_paragraph(
        "The storage subsystem comprises a NetApp FAS8200 array providing 48TB of "
        "usable NFS storage and a Pure Storage FlashArray//X50 with 12TB of block "
        "storage for database workloads. Current utilization stands at 67% for NFS "
        "and 81% for block storage, with projected growth of 15% annually."
    )

    # Figure 3
    add_figure_caption(doc, 3, "Storage utilization trends over the past 12 months by tier")

    # ---- Section 3: Target Cloud Architecture ----
    doc.add_heading("3. Target Cloud Architecture", level=1)
    doc.add_paragraph(
        "The target architecture leverages a multi-region deployment across three "
        "availability zones. Primary workloads will be hosted in the us-east-1 region "
        "with disaster recovery in eu-west-1. The architecture follows a microservices "
        "pattern with containerized application components orchestrated by Kubernetes."
    )

    doc.add_heading("3.1 Network Design", level=2)
    doc.add_paragraph(
        "The cloud network design implements a hub-and-spoke VPC topology with "
        "centralized egress through a transit gateway. Network security groups enforce "
        "least-privilege access at the subnet level, while web application firewalls "
        "protect public-facing endpoints. Direct Connect provides dedicated 10Gbps "
        "connectivity to the on-premises data center during the migration period."
    )

    # Figure 4
    add_figure_caption(doc, 4, "Target cloud VPC architecture with hub-and-spoke topology")

    doc.add_heading("3.2 Container Orchestration", level=2)
    doc.add_paragraph(
        "Application workloads are packaged as Docker containers and deployed to a "
        "managed Kubernetes cluster with auto-scaling node groups. The cluster is "
        "configured with separate node pools for compute-intensive, memory-intensive, "
        "and general-purpose workloads. Horizontal Pod Autoscaler (HPA) policies "
        "maintain response times below 200ms at the 95th percentile."
    )

    # ---- Section 4: Migration Execution Plan ----
    doc.add_heading("4. Migration Execution Plan", level=1)
    doc.add_paragraph(
        "The migration follows a phased approach over 16 weeks. Each phase includes "
        "planning, execution, testing, and validation sub-phases. A dedicated war room "
        "is established during migration weekends, staffed by infrastructure, application, "
        "and database engineering leads."
    )

    doc.add_heading("4.1 Data Migration Strategy", level=2)
    doc.add_paragraph(
        "Database migration employs a dual-write pattern during the transition period. "
        "Initial bulk data transfer uses AWS Database Migration Service (DMS) with "
        "ongoing change data capture (CDC) to maintain synchronization. The cutover "
        "window is planned for a 4-hour maintenance window with automated rollback "
        "procedures if data validation checks fail."
    )

    # Figure 5
    add_figure_caption(doc, 5, "Data migration workflow showing CDC replication and validation checkpoints")

    doc.add_heading("4.2 Application Cutover Sequence", level=2)
    doc.add_paragraph(
        "Applications are migrated in dependency order, starting with stateless "
        "frontend services and progressing to stateful backend systems. Blue-green "
        "deployment enables instant rollback at each stage. DNS-based traffic shifting "
        "provides gradual ramp-up from 10% to 100% over a 48-hour observation period."
    )

    # ---- Section 5: Post-Migration Optimization ----
    doc.add_heading("5. Post-Migration Optimization", level=1)
    doc.add_paragraph(
        "Following successful migration and a two-week stabilization period, the "
        "optimization phase focuses on right-sizing instances, implementing reserved "
        "capacity purchasing, and enabling advanced monitoring. Cost optimization "
        "targets a 20% reduction from initial cloud spend through automated scheduling "
        "and spot instance utilization for batch workloads."
    )

    # Figure 6
    add_figure_caption(doc, 6, "Projected cost savings breakdown by optimization category over 24 months")

    doc.add_heading("5.1 Monitoring and Alerting", level=2)
    doc.add_paragraph(
        "A comprehensive observability stack is deployed using Prometheus for metrics "
        "collection, Grafana for visualization, and PagerDuty for incident management. "
        "Custom dashboards provide real-time visibility into application performance, "
        "infrastructure health, and cost metrics. Alert thresholds are calibrated "
        "based on baseline measurements collected during the stabilization period."
    )

    doc.add_heading("5.2 Disaster Recovery Testing", level=2)
    doc.add_paragraph(
        "Quarterly DR tests validate recovery procedures and measure RTO/RPO against "
        "targets of 4 hours and 15 minutes respectively. Automated failover is tested "
        "monthly for critical database systems, while full-scale regional failover "
        "exercises are conducted semi-annually with participation from all engineering "
        "teams and executive stakeholders."
    )

    # ---- Appendix ----
    doc.add_heading("Appendix A: Glossary of Terms", level=1)
    doc.add_paragraph(
        "CDC - Change Data Capture: A technique for identifying and tracking changes "
        "to data in a database.\n"
        "DR - Disaster Recovery: Policies and procedures for recovering IT infrastructure "
        "after a catastrophic event.\n"
        "HPA - Horizontal Pod Autoscaler: Kubernetes component that automatically scales "
        "the number of pods based on observed metrics.\n"
        "RTO - Recovery Time Objective: Maximum acceptable downtime after a failure.\n"
        "RPO - Recovery Point Objective: Maximum acceptable data loss measured in time.\n"
        "VPC - Virtual Private Cloud: An isolated virtual network within a cloud provider."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
