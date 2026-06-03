"""
Initial Setup: Legal contract with 8 sections, no protection
Task ID: writer_legal_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    """Add a heading with contract-style formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_body(doc, text, bold=False, space_after=Pt(6)):
    """Add a body paragraph with consistent formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = space_after
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    return para


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ============================================================
    # DOCUMENT TITLE
    # ============================================================
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = subtitle.add_run('Contract No. PSA-2025-0847')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = date_para.add_run('Effective Date: March 15, 2025')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # ============================================================
    # SECTION 1: PARTIES AND RECITALS
    # ============================================================
    add_heading(doc, 'Section 1: Parties and Recitals')

    add_body(doc, 'This Professional Services Agreement ("Agreement") is entered into as of '
             'March 15, 2025 ("Effective Date"), by and between:')

    add_body(doc, 'Meridian Technology Solutions, Inc., a Delaware corporation with its '
             'principal offices located at 2400 Innovation Drive, Suite 1200, San Francisco, '
             'California 94105 (hereinafter referred to as "Provider" or "Meridian");', bold=True)

    add_body(doc, 'AND')

    add_body(doc, 'Cascade Financial Group, LLC, a Washington limited liability company with '
             'its principal offices located at 750 Pacific Avenue, Suite 800, Seattle, '
             'Washington 98101 (hereinafter referred to as "Client" or "Cascade");', bold=True)

    add_body(doc, 'WHEREAS, the Client desires to engage the Provider to perform certain '
             'professional technology consulting and implementation services as described herein; and')

    add_body(doc, 'WHEREAS, the Provider possesses the expertise, personnel, and resources '
             'necessary to deliver such services in accordance with industry best practices; and')

    add_body(doc, 'WHEREAS, both parties wish to establish the terms and conditions under which '
             'such services shall be provided;')

    add_body(doc, 'NOW, THEREFORE, in consideration of the mutual covenants and agreements '
             'contained herein, and for other good and valuable consideration, the receipt and '
             'sufficiency of which are hereby acknowledged, the parties agree as follows:')

    # ============================================================
    # SECTION 2: SCOPE OF SERVICES
    # ============================================================
    add_heading(doc, 'Section 2: Scope of Services')

    add_body(doc, '2.1 Primary Services. The Provider shall deliver the following professional '
             'services to the Client during the term of this Agreement:')

    for item in [
        'Enterprise resource planning (ERP) system assessment and gap analysis for the '
        "Client's existing Oracle and SAP infrastructure across 14 regional offices",
        'Custom software development for the Client\'s proprietary trading platform, '
        'including real-time data analytics modules, risk assessment dashboards, and '
        'automated compliance reporting tools',
        'Cloud migration services for legacy on-premise systems to AWS GovCloud, including '
        'security hardening, data encryption at rest and in transit, and disaster recovery planning',
        'Staff augmentation with a minimum of 12 senior-level consultants holding relevant '
        'certifications (AWS Solutions Architect Professional, PMP, CISSP)',
        'Quarterly technology roadmap reviews and strategic planning sessions with Client '
        'executive leadership'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_body(doc, '2.2 Service Levels. All services shall be performed in accordance with the '
             'Service Level Agreement attached hereto as Exhibit A, which establishes minimum '
             'uptime guarantees of 99.95%, maximum response times of 15 minutes for Priority 1 '
             'incidents, and quarterly performance benchmarking.')

    add_body(doc, '2.3 Change Orders. Any modifications to the scope of services described in '
             'Section 2.1 shall require a written change order signed by authorized '
             'representatives of both parties. Change orders shall specify the additional scope, '
             'timeline impact, and cost adjustments.')

    # ============================================================
    # SECTION 3: COMPENSATION AND PAYMENT TERMS
    # ============================================================
    add_heading(doc, 'Section 3: Compensation and Payment Terms')

    add_body(doc, '3.1 Base Compensation. The Client shall pay the Provider a total base fee of '
             'Four Million Two Hundred Fifty Thousand Dollars ($4,250,000.00) for services '
             'rendered under this Agreement, payable in accordance with the milestone schedule '
             'set forth below:')

    # Payment milestone table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Milestone', 'Description', 'Due Date', 'Amount']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    milestones = [
        ['M1', 'Project Kickoff & Assessment', 'April 1, 2025', '$637,500.00'],
        ['M2', 'Architecture Design Approval', 'June 15, 2025', '$850,000.00'],
        ['M3', 'Development Phase Completion', 'October 1, 2025', '$1,062,500.00'],
        ['M4', 'User Acceptance Testing', 'January 15, 2026', '$850,000.00'],
        ['M5', 'Go-Live & Hypercare Period', 'March 31, 2026', '$850,000.00'],
    ]
    for r, row_data in enumerate(milestones, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    add_body(doc, '3.2 Payment Terms. All invoices shall be payable within thirty (30) calendar '
             'days of receipt. Late payments shall accrue interest at a rate of 1.5% per month '
             'or the maximum rate permitted by applicable law, whichever is less.')

    add_body(doc, '3.3 Expenses. Reasonable and documented travel, lodging, and incidental '
             'expenses incurred by Provider personnel in connection with on-site work shall be '
             'reimbursed by the Client within forty-five (45) days of submission, subject to '
             "compliance with the Client's expense policy attached as Exhibit B.")

    # ============================================================
    # SECTION 4: CONFIDENTIALITY AND NON-DISCLOSURE
    # ============================================================
    add_heading(doc, 'Section 4: Confidentiality and Non-Disclosure')

    add_body(doc, '4.1 Definition of Confidential Information. "Confidential Information" means '
             'all non-public information disclosed by either party to the other, whether orally, '
             'in writing, or by inspection of tangible objects, that is designated as confidential '
             'or that reasonably should be understood to be confidential given the nature of the '
             'information and the circumstances of disclosure. Confidential Information includes, '
             'without limitation:')

    for item in [
        'Trade secrets, proprietary algorithms, source code, and technical specifications',
        'Business plans, financial data, customer lists, and pricing strategies',
        'Employee information, compensation data, and organizational structures',
        'Security protocols, vulnerability assessments, and penetration testing results',
        'Any information marked or identified as "Confidential," "Proprietary," or with '
        'similar designations'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_body(doc, '4.2 Obligations. Each party agrees to: (a) maintain the confidentiality of '
             'the other party\'s Confidential Information using at least the same degree of care '
             'it uses to protect its own confidential information, but in no event less than '
             'reasonable care; (b) not disclose such information to any third party without prior '
             'written consent; and (c) limit access to such information to employees and '
             'contractors who have a need to know and are bound by obligations of confidentiality '
             'at least as restrictive as those set forth herein.')

    add_body(doc, '4.3 Duration. The obligations of confidentiality shall survive the '
             'termination or expiration of this Agreement for a period of five (5) years, except '
             'with respect to trade secrets, which shall be protected for so long as they remain '
             'trade secrets under applicable law.')

    add_body(doc, '4.4 Exclusions. Confidential Information does not include information that: '
             '(a) is or becomes publicly available through no fault of the receiving party; '
             '(b) was known to the receiving party prior to disclosure; (c) is independently '
             'developed by the receiving party without use of the disclosing party\'s Confidential '
             'Information; or (d) is rightfully obtained from a third party without restriction on '
             'disclosure.')

    # ============================================================
    # SECTION 5: STANDARD TERMS AND CONDITIONS
    # ============================================================
    add_heading(doc, 'Section 5: Standard Terms and Conditions')

    add_body(doc, '5.1 Representations and Warranties. Each party represents and warrants that: '
             '(a) it has the legal power and authority to enter into this Agreement; (b) the '
             'execution of this Agreement does not conflict with any other agreement or obligation '
             'to which it is a party; and (c) it shall comply with all applicable federal, state, '
             'and local laws, regulations, and ordinances in the performance of its obligations '
             'under this Agreement.')

    add_body(doc, '5.2 Provider Warranties. The Provider further represents and warrants that: '
             '(a) all services will be performed in a professional and workmanlike manner '
             'consistent with generally accepted industry standards; (b) all deliverables will '
             'conform to the specifications set forth in the applicable statement of work; '
             '(c) Provider personnel assigned to the project will possess the qualifications, '
             'skills, and experience described in their respective profiles; and (d) the '
             'deliverables will not infringe upon the intellectual property rights of any '
             'third party.')

    add_body(doc, '5.3 Limitation of Liability. EXCEPT FOR BREACHES OF CONFIDENTIALITY '
             'OBLIGATIONS, INDEMNIFICATION OBLIGATIONS, OR WILLFUL MISCONDUCT, NEITHER PARTY '
             "SHALL BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, "
             'OR PUNITIVE DAMAGES, INCLUDING WITHOUT LIMITATION LOST PROFITS, LOST REVENUE, '
             'LOST DATA, OR BUSINESS INTERRUPTION, REGARDLESS OF THE FORM OF ACTION OR THEORY '
             'OF LIABILITY, EVEN IF SUCH PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.')

    add_body(doc, "5.4 Aggregate Liability Cap. EACH PARTY'S TOTAL CUMULATIVE LIABILITY UNDER "
             'THIS AGREEMENT SHALL NOT EXCEED TWO TIMES (2x) THE TOTAL FEES PAID OR PAYABLE '
             'UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING '
             'THE EVENT GIVING RISE TO THE CLAIM. THIS LIMITATION SHALL NOT APPLY TO: (a) '
             'BREACHES OF SECTION 4 (CONFIDENTIALITY); (b) INDEMNIFICATION OBLIGATIONS UNDER '
             'SECTION 5.6; OR (c) DAMAGES ARISING FROM WILLFUL MISCONDUCT OR GROSS NEGLIGENCE.')

    add_body(doc, '5.5 Indemnification. Each party ("Indemnifying Party") shall indemnify, '
             'defend, and hold harmless the other party and its officers, directors, employees, '
             'and agents ("Indemnified Party") from and against any and all claims, damages, '
             'losses, liabilities, costs, and expenses (including reasonable attorneys\' fees) '
             'arising out of or relating to: (a) any breach of this Agreement by the Indemnifying '
             'Party; (b) any negligent or wrongful act or omission of the Indemnifying Party; or '
             '(c) any third-party claim alleging that deliverables provided by the Provider '
             'infringe upon or misappropriate any intellectual property right.')

    add_body(doc, '5.6 Insurance Requirements. The Provider shall maintain, at its own expense, '
             'the following insurance coverage throughout the term of this Agreement and for a '
             'period of two (2) years following termination:')

    for item in [
        'Commercial General Liability: $5,000,000 per occurrence / $10,000,000 aggregate',
        'Professional Liability (Errors & Omissions): $10,000,000 per claim / $20,000,000 aggregate',
        'Cyber Liability and Technology Errors & Omissions: $5,000,000 per occurrence',
        "Workers' Compensation: As required by applicable state law",
        'Commercial Automobile Liability: $1,000,000 combined single limit'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_body(doc, '5.7 Force Majeure. Neither party shall be liable for any failure or delay in '
             'performing its obligations under this Agreement to the extent that such failure or '
             'delay results from circumstances beyond the reasonable control of such party, '
             'including but not limited to: acts of God, natural disasters, epidemics or '
             'pandemics, war, terrorism, civil unrest, government actions, labor disputes, '
             'power failures, internet or telecommunications failures, or cyberattacks. The '
             'affected party shall provide prompt written notice to the other party of the force '
             'majeure event and shall use commercially reasonable efforts to mitigate its effects.')

    add_body(doc, '5.8 Compliance with Laws. Both parties shall comply with all applicable '
             'federal, state, and local laws, rules, and regulations, including but not limited '
             'to: the Sarbanes-Oxley Act, the Health Insurance Portability and Accountability '
             'Act (HIPAA), the California Consumer Privacy Act (CCPA), the General Data Protection '
             'Regulation (GDPR) to the extent applicable, the Federal Acquisition Regulation (FAR) '
             'if applicable, and all export control laws and regulations.')

    add_body(doc, '5.9 Anti-Corruption. Each party represents and warrants that it has not, and '
             'covenants that it will not, directly or indirectly, offer, promise, give, or '
             'authorize the giving of money or anything of value to any government official, '
             'political party, or candidate for political office for the purpose of influencing '
             'any act or decision of such official, party, or candidate, or to obtain an improper '
             'advantage in connection with this Agreement, in violation of the U.S. Foreign Corrupt '
             'Practices Act, the UK Bribery Act, or any other applicable anti-corruption law.')

    add_body(doc, '5.10 Assignment. Neither party may assign or transfer this Agreement, or any '
             'rights or obligations hereunder, without the prior written consent of the other '
             'party, which consent shall not be unreasonably withheld. Notwithstanding the '
             'foregoing, either party may assign this Agreement without consent in connection with '
             'a merger, acquisition, or sale of all or substantially all of its assets, provided '
             'that the assignee agrees in writing to be bound by the terms of this Agreement.')

    add_body(doc, '5.11 Notices. All notices, requests, demands, and other communications under '
             'this Agreement shall be in writing and shall be deemed duly given: (a) when '
             'delivered personally; (b) when sent by confirmed electronic mail; (c) one (1) '
             'business day after being sent by nationally recognized overnight courier; or '
             '(d) three (3) business days after being mailed by certified or registered mail, '
             'return receipt requested, postage prepaid, to the addresses set forth in this '
             'Agreement or to such other address as either party may designate in writing.')

    add_body(doc, '5.12 Governing Law and Jurisdiction. This Agreement shall be governed by and '
             'construed in accordance with the laws of the State of Delaware, without regard to '
             'its conflict of laws principles. Any dispute arising out of or relating to this '
             'Agreement shall be subject to the exclusive jurisdiction of the federal and state '
             'courts located in Wilmington, Delaware, and each party hereby consents to the '
             'personal jurisdiction of such courts.')

    # ============================================================
    # SECTION 6: INTELLECTUAL PROPERTY RIGHTS
    # ============================================================
    add_heading(doc, 'Section 6: Intellectual Property Rights')

    add_body(doc, '6.1 Client Ownership. All deliverables, work product, and materials '
             'specifically created by the Provider for the Client under this Agreement '
             '("Client IP") shall be the sole and exclusive property of the Client. The Provider '
             'hereby assigns to the Client all right, title, and interest in and to the Client '
             'IP, including all intellectual property rights therein.')

    add_body(doc, '6.2 Provider Pre-Existing IP. The Provider retains all right, title, and '
             'interest in and to its pre-existing intellectual property, including tools, '
             'frameworks, methodologies, templates, and know-how that existed prior to this '
             'Agreement or were developed independently of this Agreement ("Provider IP"). To '
             'the extent any Provider IP is incorporated into deliverables, the Provider grants '
             'the Client a perpetual, irrevocable, worldwide, royalty-free, non-exclusive license '
             'to use, modify, and sublicense such Provider IP solely in connection with the '
             "Client's use of the deliverables.")

    add_body(doc, '6.3 Open Source Components. The Provider shall disclose all open source '
             'components incorporated into deliverables and ensure compliance with applicable '
             'open source licenses. No deliverable shall incorporate any open source component '
             'subject to a "copyleft" license (e.g., GPL, AGPL) without the prior written '
             'approval of the Client.')

    # ============================================================
    # SECTION 7: TERMINATION AND DISPUTE RESOLUTION
    # ============================================================
    add_heading(doc, 'Section 7: Termination and Dispute Resolution')

    add_body(doc, '7.1 Term. This Agreement shall commence on the Effective Date and shall '
             'continue for a period of twenty-four (24) months, unless earlier terminated in '
             'accordance with this Section 7.')

    add_body(doc, '7.2 Termination for Convenience. Either party may terminate this Agreement '
             'for any reason upon ninety (90) days\' prior written notice to the other party. In '
             'the event of termination for convenience by the Client, the Client shall pay the '
             'Provider for all services rendered and expenses incurred through the effective date '
             'of termination, plus a termination fee equal to fifteen percent (15%) of the '
             'remaining unpaid balance.')

    add_body(doc, '7.3 Termination for Cause. Either party may terminate this Agreement '
             'immediately upon written notice if the other party: (a) materially breaches this '
             'Agreement and fails to cure such breach within thirty (30) days after receipt of '
             'written notice specifying the breach; (b) becomes insolvent, files for bankruptcy, '
             'or has a receiver appointed for its assets; or (c) is found to have made a material '
             'misrepresentation in connection with this Agreement.')

    add_body(doc, '7.4 Dispute Resolution. Any dispute, controversy, or claim arising out of or '
             'relating to this Agreement shall be resolved as follows:')

    for item in [
        'Step 1 - Executive Negotiation: The parties shall first attempt to resolve the '
        'dispute through good-faith negotiations between senior executives within thirty (30) days',
        'Step 2 - Mediation: If negotiation fails, the parties shall submit the dispute to '
        'non-binding mediation administered by JAMS under its Mediation Rules within sixty (60) days',
        'Step 3 - Arbitration: If mediation fails, the dispute shall be resolved by binding '
        'arbitration administered by JAMS under its Comprehensive Arbitration Rules, with three '
        'arbitrators, in Wilmington, Delaware'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_body(doc, '7.5 Survival. Sections 4 (Confidentiality), 5.3-5.5 (Liability and '
             'Indemnification), 6 (Intellectual Property), and this Section 7 shall survive '
             'the termination or expiration of this Agreement.')

    # ============================================================
    # SECTION 8: SIGNATURES AND EXECUTION
    # ============================================================
    add_heading(doc, 'Section 8: Signatures and Execution')

    add_body(doc, '8.1 Entire Agreement. This Agreement, including all exhibits and schedules '
             'attached hereto, constitutes the entire agreement between the parties with respect '
             'to the subject matter hereof and supersedes all prior and contemporaneous '
             'agreements, understandings, negotiations, and discussions, whether oral or written.')

    add_body(doc, '8.2 Amendments. This Agreement may not be amended or modified except by a '
             'written instrument executed by authorized representatives of both parties.')

    add_body(doc, '8.3 Counterparts. This Agreement may be executed in counterparts, each of '
             'which shall be deemed an original, but all of which together shall constitute one '
             'and the same instrument. Electronic signatures shall be deemed valid and binding.')

    add_body(doc, '8.4 Severability. If any provision of this Agreement is held to be invalid, '
             'illegal, or unenforceable, the remaining provisions shall continue in full force '
             'and effect.')

    doc.add_paragraph()  # spacer

    add_body(doc, 'IN WITNESS WHEREOF, the parties have executed this Agreement as of the '
             'Effective Date first written above.', bold=True)

    doc.add_paragraph()

    # Signature blocks
    add_body(doc, 'MERIDIAN TECHNOLOGY SOLUTIONS, INC.', bold=True)
    doc.add_paragraph()
    add_body(doc, '___________________________________')
    add_body(doc, 'Name: Dr. Alexandra Petrov')
    add_body(doc, 'Title: Chief Executive Officer')
    add_body(doc, 'Date: March 15, 2025')

    doc.add_paragraph()

    add_body(doc, 'CASCADE FINANCIAL GROUP, LLC', bold=True)
    doc.add_paragraph()
    add_body(doc, '___________________________________')
    add_body(doc, 'Name: Robert J. Hawthorne III')
    add_body(doc, 'Title: Managing Partner')
    add_body(doc, 'Date: March 15, 2025')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
