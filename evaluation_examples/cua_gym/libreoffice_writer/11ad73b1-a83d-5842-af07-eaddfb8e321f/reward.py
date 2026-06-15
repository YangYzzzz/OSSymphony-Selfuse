"""
Reward Script: Add light gray (#E0E0E0) background to entire TOC area
Task ID: writer_mt_074
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): TOC heading paragraph has gray shading
  Component 2 (0.4): All 25 TOC entry paragraphs have gray shading
  Component 3 (0.2): Body text outside TOC retains no gray shading
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_074'
TARGET_COLOR = 'E0E0E0'
# Perceptual tolerance for color matching (hex distance)
COLOR_TOLERANCE = 30


def hex_to_rgb(h):
    """Convert hex string like 'E0E0E0' to (r, g, b) tuple."""
    h = h.strip('#').upper()
    return (int(h[0:2], 16), int(h[1:4][:2], 16), int(h[4:6], 16))


def color_close(hex1, hex2, tolerance=COLOR_TOLERANCE):
    """Check if two hex colors are perceptually close enough."""
    try:
        r1, g1, b1 = hex_to_rgb(hex1)
        r2, g2, b2 = hex_to_rgb(hex2)
        dist = sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)
        return dist <= tolerance
    except Exception:
        return False


def get_para_shading_fill(para_element, ns):
    """Get the fill color from paragraph shading, or None."""
    pPr = para_element.find('w:pPr', ns)
    if pPr is None:
        return None
    shd = pPr.find('w:shd', ns)
    if shd is None:
        return None
    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
    return fill


def find_toc_range(doc):
    """
    Find the TOC paragraph range. The TOC starts with a 'Table of Contents'
    heading and ends before the first non-TOC-entry paragraph (body heading/content).
    Returns (start_idx, end_idx) inclusive, or (None, None) if not found.
    """
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        style_name = p.style.name if p.style else ''
        if start is None:
            if 'table of contents' in text or 'toc' == text:
                start = i
        else:
            # TOC entries are Normal style with tab-separated page numbers
            # or continuation of the TOC heading. Stop when we hit a body
            # heading (Heading 1/2/3 that is NOT the TOC title).
            if style_name.startswith('Heading') and 'table of contents' not in text:
                end = i - 1
                break
            # Also stop if we hit substantial body text (non-empty, no tab)
            # But skip empty paragraphs within TOC area
    if start is not None and end is None:
        # Scan backwards from end to find last TOC entry
        end = len(doc.paragraphs) - 1
    return start, end


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Find the TOC range
    toc_start, toc_end = find_toc_range(doc)
    if toc_start is None:
        print("FAIL: Could not locate 'Table of Contents' heading in document")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: TOC range detected: paragraphs [{toc_start}] to [{toc_end}]")
    print(f"INFO: TOC heading: {doc.paragraphs[toc_start].text!r}")

    # Component 1: TOC heading paragraph has gray shading (0.4 points)
    try:
        toc_heading = doc.paragraphs[toc_start]
        fill = get_para_shading_fill(toc_heading._element, ns)
        if fill and color_close(fill, TARGET_COLOR):
            print(f"PASS: Component 1 — TOC heading has shading fill={fill} (close to #{TARGET_COLOR}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — TOC heading shading fill={fill}, expected close to #{TARGET_COLOR}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All TOC entry paragraphs have gray shading (0.4 points)
    # TOC entries are paragraphs from toc_start+1 to toc_end
    try:
        toc_entries_start = toc_start + 1
        toc_entries_end = toc_end
        total_entries = toc_entries_end - toc_entries_start + 1

        if total_entries <= 0:
            print("FAIL: Component 2 — No TOC entries found after heading")
        else:
            shaded_count = 0
            for idx in range(toc_entries_start, toc_entries_end + 1):
                p = doc.paragraphs[idx]
                # Skip empty paragraphs
                if not p.text.strip():
                    total_entries -= 1
                    continue
                fill = get_para_shading_fill(p._element, ns)
                if fill and color_close(fill, TARGET_COLOR):
                    shaded_count += 1
                else:
                    print(f"  DETAIL: Para [{idx}] missing shading (fill={fill}): {p.text[:50]!r}")

            if total_entries > 0:
                ratio = shaded_count / total_entries
                comp2_score = round(0.4 * ratio, 2)
                total_score += comp2_score
                print(f"{'PASS' if ratio == 1.0 else 'PARTIAL'}: Component 2 — {shaded_count}/{total_entries} TOC entries have gray shading ({comp2_score} pts)")
            else:
                print("FAIL: Component 2 — No non-empty TOC entries found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: TOC has gray shading AND body text outside TOC does NOT (0.2 points)
    # This is a compound check: TOC must be shaded (task change) AND body must be unshaded.
    # This ensures the component only passes when the task has been completed (TOC shaded)
    # while body remains clean. Fails on initial_env because TOC isn't shaded.
    try:
        # First gate: TOC heading must have shading (otherwise task not done)
        toc_heading_fill = get_para_shading_fill(doc.paragraphs[toc_start]._element, ns)
        toc_is_shaded = toc_heading_fill is not None and color_close(toc_heading_fill, TARGET_COLOR)

        if not toc_is_shaded:
            print("FAIL: Component 3 — TOC heading not shaded, so compound check fails")
        else:
            body_start = toc_end + 1
            body_paras_checked = 0
            body_paras_wrongly_shaded = 0
            # Check up to 20 body paragraphs after TOC
            for idx in range(body_start, min(body_start + 20, len(doc.paragraphs))):
                p = doc.paragraphs[idx]
                if not p.text.strip():
                    continue
                body_paras_checked += 1
                fill = get_para_shading_fill(p._element, ns)
                if fill and color_close(fill, TARGET_COLOR):
                    body_paras_wrongly_shaded += 1
                    print(f"  DETAIL: Body para [{idx}] wrongly has gray shading: {p.text[:50]!r}")

            if body_paras_checked == 0:
                print("WARN: Component 3 — No body paragraphs found after TOC to check")
                total_score += 0.2
                print(f"PASS: Component 3 — TOC shaded + no body paragraphs to check (0.2 pts)")
            elif body_paras_wrongly_shaded == 0:
                print(f"PASS: Component 3 — TOC shaded + {body_paras_checked} body paragraphs unshaded (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — {body_paras_wrongly_shaded}/{body_paras_checked} body paragraphs wrongly have gray shading")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
