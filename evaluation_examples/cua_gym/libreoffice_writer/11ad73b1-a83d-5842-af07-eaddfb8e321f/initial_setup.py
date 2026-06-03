"""
Initial Setup: Create Training_Guide.docx with a TOC and body content (no shading).
Task ID: writer_mt_074
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_entry(doc, text, level, page_num):
    """Add a fake TOC entry paragraph styled like a real TOC."""
    para = doc.add_paragraph()
    # Set TOC style indent based on level
    if level == 1:
        para.paragraph_format.left_indent = Inches(0)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
    elif level == 2:
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(text)
        run.font.size = Pt(11)
    else:
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run(text)
        run.font.size = Pt(10)

    # Add tab + page number
    tab_run = para.add_run('\t')
    pg_run = para.add_run(str(page_num))
    pg_run.font.size = run.font.size

    # Add a right-aligned tab stop with dot leader
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    para.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    return para


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Employee Training Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = subtitle.add_run('Meridian Solutions Inc. — 2025 Edition')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # --- Table of Contents Header ---
    toc_header = doc.add_heading('Table of Contents', level=1)

    # --- TOC Entries (25 entries, 3 levels) ---
    toc_entries = [
        # (text, level, page)
        ('1. Introduction to the Company', 1, 3),
        ('1.1 Company History', 2, 3),
        ('1.2 Mission and Vision', 2, 4),
        ('1.2.1 Core Values', 3, 4),
        ('1.2.2 Strategic Objectives', 3, 5),
        ('2. Onboarding Process', 1, 6),
        ('2.1 First Day Checklist', 2, 6),
        ('2.2 HR Documentation', 2, 7),
        ('2.2.1 Benefits Enrollment', 3, 7),
        ('2.2.2 Tax Forms and Payroll', 3, 8),
        ('2.3 IT Setup and Access', 2, 9),
        ('3. Safety and Compliance', 1, 10),
        ('3.1 Workplace Safety Protocols', 2, 10),
        ('3.1.1 Emergency Evacuation Procedures', 3, 11),
        ('3.1.2 Fire Safety Equipment', 3, 11),
        ('3.2 Data Security Policies', 2, 12),
        ('3.3 Anti-Harassment Training', 2, 13),
        ('4. Professional Development', 1, 14),
        ('4.1 Skills Assessment Framework', 2, 14),
        ('4.2 Training Programs Available', 2, 15),
        ('4.2.1 Technical Certifications', 3, 15),
        ('4.2.2 Leadership Development', 3, 16),
        ('4.3 Mentorship Program', 2, 17),
        ('5. Performance Evaluation', 1, 18),
        ('5.1 Quarterly Review Process', 2, 18),
    ]

    for text, level, page in toc_entries:
        add_toc_entry(doc, text, level, page)

    # --- Page break after TOC ---
    doc.add_page_break()

    # --- Body Content ---
    doc.add_heading('1. Introduction to the Company', level=1)
    doc.add_paragraph(
        'Meridian Solutions Inc. was founded in 2003 by Elena Rodriguez and '
        'James Park with a vision to transform enterprise software consulting. '
        'Over the past two decades, the company has grown from a small team of '
        'five to over 2,400 employees across 12 global offices.'
    )
    doc.add_paragraph(
        'Our headquarters in Austin, Texas serves as the hub for innovation '
        'and strategic planning. We maintain regional offices in London, '
        'Singapore, São Paulo, and Toronto to support our international client base.'
    )

    doc.add_heading('1.1 Company History', level=2)
    doc.add_paragraph(
        'The company began as a boutique consulting firm specializing in '
        'database optimization for healthcare providers. In 2008, under the '
        'leadership of CTO Dr. Anika Patel, Meridian expanded into cloud '
        'infrastructure services, which now accounts for 45% of annual revenue.'
    )

    doc.add_heading('1.2 Mission and Vision', level=2)
    doc.add_paragraph(
        'Our mission is to empower organizations through intelligent technology '
        'solutions that drive measurable business outcomes. We believe in '
        'sustainable growth, ethical practices, and continuous improvement.'
    )

    doc.add_heading('2. Onboarding Process', level=1)
    doc.add_paragraph(
        'All new hires at Meridian Solutions complete a structured two-week '
        'onboarding program. During the first week, employees meet with their '
        'direct manager, department head, and assigned mentor. The second week '
        'focuses on role-specific training and system access provisioning.'
    )

    doc.add_heading('2.1 First Day Checklist', level=2)
    doc.add_paragraph('On your first day, please bring the following:', style='List Bullet')
    doc.add_paragraph('Government-issued photo ID', style='List Bullet')
    doc.add_paragraph('Signed offer letter', style='List Bullet')
    doc.add_paragraph('Banking information for direct deposit', style='List Bullet')
    doc.add_paragraph('Emergency contact details', style='List Bullet')

    doc.add_heading('3. Safety and Compliance', level=1)
    doc.add_paragraph(
        'Meridian Solutions maintains rigorous safety and compliance standards '
        'in accordance with OSHA regulations and industry best practices. '
        'All employees are required to complete annual safety certification '
        'training, which includes both online modules and in-person workshops.'
    )

    doc.add_heading('4. Professional Development', level=1)
    doc.add_paragraph(
        'We invest heavily in our employees\' professional growth. Each team '
        'member receives an annual training budget of $3,500 and is encouraged '
        'to pursue relevant certifications. Our internal learning platform, '
        'MeridianLearn, offers over 500 courses across technical and soft skills.'
    )

    doc.add_heading('5. Performance Evaluation', level=1)
    doc.add_paragraph(
        'Performance reviews at Meridian follow a quarterly cadence. Each cycle '
        'includes self-assessment, peer feedback, and a one-on-one review with '
        'your manager. Goals are set collaboratively using the OKR framework, '
        'and progress is tracked through our internal performance dashboard.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
