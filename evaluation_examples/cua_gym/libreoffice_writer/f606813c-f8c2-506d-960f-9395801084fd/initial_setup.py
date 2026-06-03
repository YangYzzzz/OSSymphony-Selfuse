"""
Initial Setup: Recruitment tracking document with intro paragraph only (no table).
Task ID: writer_hr_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title heading
    heading = doc.add_heading('Recruitment Tracker Q1 2026', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction paragraph
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'This document tracks the recruitment pipeline for Q1 2026 across all departments. '
        'It provides an overview of open positions, application volumes, interview scheduling, '
        'and offer statuses to support the HR team in managing the hiring workflow efficiently.'
    )
    intro_run.font.size = Pt(11)
    intro_run.font.name = 'Calibri'

    # Add a blank paragraph for spacing
    doc.add_paragraph()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
