"""
Reward Script: Recruitment tracking table with color-coded Status column
Task ID: writer_hr_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table exists with 9 rows x 7 columns
  Component 2 (0.3): Correct headers and all 8 data rows populated
  Component 3 (0.4): Status column color-coded (green=Filled, yellow=In Progress, red=On Hold)
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_050'

# Expected headers
EXPECTED_HEADERS = ['Position', 'Department', 'Posting Date',
                    'Applications Received', 'Interviews Scheduled',
                    'Offers Made', 'Status']

# Color mappings for status (using perceptual distance tolerance)
# Green variants for "Filled"
GREEN_TARGETS = [(0x00, 0xB0, 0x50), (0x00, 0xFF, 0x00), (0x00, 0x80, 0x00),
                 (0x92, 0xD0, 0x50), (0x00, 0xB0, 0x00)]
# Yellow variants for "In Progress"
YELLOW_TARGETS = [(0xFF, 0xD9, 0x66), (0xFF, 0xFF, 0x00), (0xFF, 0xC0, 0x00),
                  (0xBF, 0xBF, 0x00), (0xFF, 0xE6, 0x99)]
# Red variants for "On Hold"
RED_TARGETS = [(0xFF, 0x4B, 0x4B), (0xFF, 0x00, 0x00), (0xC0, 0x00, 0x00),
               (0xFF, 0x66, 0x66), (0xFF, 0x33, 0x33)]


def color_distance(c1, c2):
    """Euclidean RGB distance."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def is_color_close(fill_hex, target_list, threshold=120):
    """Check if a hex color string is close to any target RGB tuple."""
    if not fill_hex or fill_hex.lower() in ('auto', 'none', ''):
        return False
    try:
        fill_hex = fill_hex.lstrip('#')
        r = int(fill_hex[0:2], 16)
        g = int(fill_hex[2:4], 16)
        b = int(fill_hex[4:6], 16)
        for target in target_list:
            if color_distance((r, g, b), target) < threshold:
                return True
    except (ValueError, IndexError):
        pass
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions (0.3 points)
    # Initial env has no tables, golden has a 9x7 table
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 -- No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 9 and num_cols == 7:
                print(f"PASS: Component 1 -- Table found with {num_rows} rows x {num_cols} cols (0.3 pts)")
                total_score += 0.3
            elif num_rows >= 2 and num_cols == 7:
                # Partial: correct columns but wrong row count
                if num_rows >= 2:
                    print(f"PARTIAL: Component 1 -- Table has {num_rows} rows x {num_cols} cols (expected 9x7) (0.15 pts)")
                    total_score += 0.15
            else:
                print(f"FAIL: Component 1 -- Table has {num_rows} rows x {num_cols} cols (expected 9x7)")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Gate: Need at least one table to continue
    if len(doc.tables) == 0:
        print("\nScore: 0.0/1.0")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    from docx.oxml.ns import qn

    # Component 2: Correct headers and all data rows populated (0.3 points)
    # Initial env has no table at all, so this only passes on golden
    try:
        # Check headers (0.15)
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        headers_match = 0
        for expected in EXPECTED_HEADERS:
            if any(expected.lower() in h.lower() for h in header_cells):
                headers_match += 1
        header_ratio = headers_match / len(EXPECTED_HEADERS)

        # Check data rows populated (0.15)
        populated_rows = 0
        for ri in range(1, len(table.rows)):
            row_texts = [cell.text.strip() for cell in table.rows[ri].cells]
            # A row is populated if at least 5 of 7 cells have content
            if sum(1 for t in row_texts if t) >= 5:
                populated_rows += 1

        comp2_score = 0.0
        if header_ratio >= 0.85:
            comp2_score += 0.15
            print(f"PASS: Component 2a -- Headers match ({headers_match}/{len(EXPECTED_HEADERS)})")
        else:
            print(f"FAIL: Component 2a -- Headers match only {headers_match}/{len(EXPECTED_HEADERS)}: {header_cells}")

        if populated_rows >= 8:
            comp2_score += 0.15
            print(f"PASS: Component 2b -- {populated_rows} data rows populated (0.15 pts)")
        elif populated_rows >= 4:
            partial = 0.075
            comp2_score += partial
            print(f"PARTIAL: Component 2b -- Only {populated_rows}/8 data rows populated ({partial} pts)")
        else:
            print(f"FAIL: Component 2b -- Only {populated_rows}/8 data rows populated")

        if comp2_score > 0:
            total_score += comp2_score
        print(f"Component 2 total: {comp2_score} pts")

    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Status column color-coded (0.4 points)
    # Check that status cells have appropriate background shading
    # Initial env has no table, so this only passes on golden
    try:
        status_col_idx = 6  # Last column = Status
        correct_colors = 0
        total_data_rows = 0

        for ri in range(1, len(table.rows)):
            cell = table.rows[ri].cells[status_col_idx]
            status_text = cell.text.strip().lower()
            if not status_text:
                continue

            total_data_rows += 1

            # Get cell shading
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            shd = None
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))

            if shd is None:
                print(f"FAIL: Row {ri} status='{status_text}' has no shading")
                continue

            fill = shd.get(qn('w:fill'))

            if 'filled' in status_text:
                if is_color_close(fill, GREEN_TARGETS):
                    correct_colors += 1
                    print(f"PASS: Row {ri} status='Filled' has green fill ({fill})")
                else:
                    print(f"FAIL: Row {ri} status='Filled' fill={fill} is not green")
            elif 'in progress' in status_text:
                if is_color_close(fill, YELLOW_TARGETS):
                    correct_colors += 1
                    print(f"PASS: Row {ri} status='In Progress' has yellow fill ({fill})")
                else:
                    print(f"FAIL: Row {ri} status='In Progress' fill={fill} is not yellow")
            elif 'on hold' in status_text:
                if is_color_close(fill, RED_TARGETS):
                    correct_colors += 1
                    print(f"PASS: Row {ri} status='On Hold' has red fill ({fill})")
                else:
                    print(f"FAIL: Row {ri} status='On Hold' fill={fill} is not red")
            else:
                print(f"INFO: Row {ri} status='{status_text}' -- unknown status, skipping color check")

        if total_data_rows > 0:
            color_ratio = correct_colors / total_data_rows
            comp3_score = round(0.4 * color_ratio, 2)
            print(f"Component 3: {correct_colors}/{total_data_rows} status cells correctly colored ({comp3_score} pts)")
            if comp3_score > 0:
                total_score += comp3_score
        else:
            print("FAIL: Component 3 -- No data rows with status text found")

    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
