"""
Reward Script: Emergency Contact Form Template
Task ID: writer_hr_048
Domain: libreoffice_writer
Scoring:
  Component 1: Title heading "Emergency Contact Form" (0.15)
  Component 2: Employee info table with correct fields (0.25)
  Component 3: "Emergency Contacts" section heading (0.10)
  Component 4: Emergency contacts table structure & headers (0.30)
  Component 5: Annual update note at bottom (0.20)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_048'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify emergency contact form template creation.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_para_texts = [p.text.strip() for p in doc.paragraphs]
    all_para_lower = [t.lower() for t in all_para_texts]

    # Component 1: Title heading "Emergency Contact Form" (0.15 points)
    try:
        has_title = False
        for p in doc.paragraphs:
            if 'emergency contact form' in p.text.strip().lower():
                # Check it's a heading style
                if 'heading' in (p.style.name or '').lower() or 'title' in (p.style.name or '').lower():
                    has_title = True
                    break
                # Also accept if it's simply present with bold formatting
                elif any(r.bold for r in p.runs if r.text.strip()):
                    has_title = True
                    break
        if has_title:
            print(f"PASS: Component 1 — Title 'Emergency Contact Form' found as heading (0.15 pts)")
            total_score += 0.15
        else:
            # Check if text exists at all even without heading style
            found_text = any('emergency contact form' in t for t in all_para_lower)
            if found_text:
                print(f"PARTIAL: Component 1 — Title text found but not as heading style (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 1 — Title 'Emergency Contact Form' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Employee info table with correct field labels (0.25 points)
    try:
        emp_table_found = False
        emp_fields_score = 0.0
        expected_fields = ['employee name', 'employee id', 'department']

        for table in doc.tables:
            # Look for a table with 2 columns containing employee info fields
            if len(table.columns) == 2:
                field_labels = []
                for row in table.rows:
                    cell_text = row.cells[0].text.strip().lower()
                    field_labels.append(cell_text)

                matched = sum(1 for ef in expected_fields if any(ef in fl for fl in field_labels))
                if matched >= 2:
                    emp_table_found = True
                    emp_fields_score = matched / len(expected_fields)
                    break

        if emp_table_found:
            pts = round(0.25 * emp_fields_score, 2)
            print(f"PASS: Component 2 — Employee info table found, {int(emp_fields_score*3)}/3 fields matched ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 2 — Employee info table with correct fields not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: "Emergency Contacts" section heading (0.10 points)
    try:
        has_ec_heading = False
        for p in doc.paragraphs:
            if 'emergency contact' in p.text.strip().lower() and 'form' not in p.text.strip().lower():
                if 'heading' in (p.style.name or '').lower() or 'title' in (p.style.name or '').lower():
                    has_ec_heading = True
                    break
                elif any(r.bold for r in p.runs if r.text.strip()):
                    has_ec_heading = True
                    break
        if has_ec_heading:
            print(f"PASS: Component 3 — 'Emergency Contacts' section heading found (0.10 pts)")
            total_score += 0.10
        else:
            # Accept even without heading style if the text is present
            found_ec = any('emergency contact' in t and 'form' not in t for t in all_para_lower)
            if found_ec:
                print(f"PARTIAL: Component 3 — 'Emergency Contacts' text found but not as heading (0.03 pts)")
                total_score += 0.03
            else:
                print(f"FAIL: Component 3 — 'Emergency Contacts' section heading not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Emergency contacts table — 4+ rows, 5 columns, correct headers (0.30 points)
    try:
        ec_table_found = False
        ec_score = 0.0
        expected_headers = ['contact name', 'relationship', 'phone number', 'email', 'address']

        for table in doc.tables:
            if len(table.columns) >= 5:
                # Check header row
                header_cells = [table.cell(0, c).text.strip().lower() for c in range(min(5, len(table.columns)))]
                header_matches = sum(1 for eh in expected_headers if any(eh in hc for hc in header_cells))

                if header_matches >= 3:
                    ec_table_found = True
                    # Score: header correctness (60%) + row count (40%)
                    header_pct = header_matches / len(expected_headers)
                    row_count = len(table.rows)
                    # Need at least 4 rows (1 header + 3 data)
                    row_pct = min(1.0, (row_count - 1) / 3.0) if row_count >= 2 else 0.0
                    ec_score = 0.6 * header_pct + 0.4 * row_pct
                    break

        if ec_table_found:
            pts = round(0.30 * ec_score, 2)
            print(f"PASS: Component 4 — Emergency contacts table found, score={ec_score:.2f} ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 4 — Emergency contacts table (5 cols, correct headers) not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Annual update note at bottom (0.20 points)
    try:
        has_annual_note = False
        # Check the last few non-empty paragraphs for annual/update keywords
        non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
        if non_empty_paras:
            # Check last 3 paragraphs
            for p in non_empty_paras[-3:]:
                text_lower = p.text.strip().lower()
                if ('annual' in text_lower or 'yearly' in text_lower) and ('update' in text_lower or 'review' in text_lower):
                    has_annual_note = True
                    break

        if has_annual_note:
            print(f"PASS: Component 5 — Annual update note found at bottom (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 5 — Annual update note not found in last paragraphs")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
