"""
Reward Script: Bibliography entries and formatted bibliography section in a Writer document
Task ID: writer_bs_015
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): "Bibliography" heading exists at the end of the document
  Component 2 (0.50): All 5 required bibliography entries present with correct author/journal/year
  Component 3 (0.20): Entries are sorted alphabetically by author
  Component 4 (0.15): Exactly 5 entries (no extra, no missing)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_015'

# Expected bibliography entries: (author_last_name, journal_keyword, year)
EXPECTED_ENTRIES = [
    ("Adams", "Nature", "2019"),
    ("Baker", "Science", "2020"),
    ("Clark", "Cell", "2018"),
    ("Davis", "PNAS", "2021"),  # Also accept "Proceedings of the National Academy"
    ("Evans", "JACS", "2017"),  # Also accept "Journal of the American Chemical Society"
]

# Alphabetical order of author last names
EXPECTED_ORDER = ["Adams", "Baker", "Clark", "Davis", "Evans"]


def find_bibliography_section(doc):
    """
    Find the bibliography heading and return the index of the heading paragraph.
    Returns (heading_index, heading_text) or (None, None) if not found.
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if 'bibliography' in text or 'references' in text:
            style = para.style.name if para.style else ''
            # Accept heading styles or bold text
            if 'heading' in style.lower():
                return (i, para.text.strip())
            # Also check if it's a bold paragraph acting as heading
            if para.runs and all(r.bold for r in para.runs if r.text.strip()):
                return (i, para.text.strip())
    return (None, None)


def extract_bib_entries(doc, heading_index):
    """
    Extract bibliography entry paragraphs after the heading.
    Returns list of paragraph texts.
    """
    entries = []
    for i in range(heading_index + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text:  # skip empty paragraphs
            # Check if this looks like a bibliography entry (has author info)
            entries.append(text)
    return entries


def check_entry_matches(entry_text, author, journal, year):
    """
    Check if a bibliography entry text contains the expected author, journal, and year.
    """
    text_lower = entry_text.lower()
    author_found = author.lower() in text_lower

    # Handle journal name variants
    journal_found = False
    if journal == "PNAS":
        journal_found = ("pnas" in text_lower or
                         "proceedings of the national academy" in text_lower or
                         "proc. natl. acad. sci." in text_lower)
    elif journal == "JACS":
        journal_found = ("jacs" in text_lower or
                         "journal of the american chemical society" in text_lower or
                         "j. am. chem. soc." in text_lower)
    else:
        journal_found = journal.lower() in text_lower

    year_found = year in entry_text

    return author_found and journal_found and year_found


def get_author_from_entry(entry_text):
    """
    Extract the first author last name from a bibliography entry.
    Tries common citation formats.
    """
    # Try patterns like "[1] Adams, ..." or "Adams, ..." or "1. Adams, ..."
    patterns = [
        r'^\[?\d+\]?\s*([A-Z][a-z]+)',   # [1] Adams or 1] Adams
        r'^\d+\.\s*([A-Z][a-z]+)',         # 1. Adams
        r'^([A-Z][a-z]+)',                  # Adams at start
    ]
    for pat in patterns:
        m = re.match(pat, entry_text.strip())
        if m:
            return m.group(1)
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: "Bibliography" heading exists at end of document (0.15 points)
    try:
        heading_idx, heading_text = find_bibliography_section(doc)
        if heading_idx is not None:
            # Verify it's near the end (after most content paragraphs)
            total_paras = len(doc.paragraphs)
            # The heading should be in the last ~25% of paragraphs
            if heading_idx >= total_paras * 0.5:
                print(f"PASS: Component 1 - Bibliography heading '{heading_text}' found at P{heading_idx} (near end) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 - Bibliography heading found at P{heading_idx} but not near end of document (total: {total_paras})")
        else:
            print("FAIL: Component 1 - No Bibliography/References heading found in document")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: All 5 required bibliography entries present with correct content (0.50 points)
    # Each entry worth 0.10 points
    try:
        if heading_idx is not None:
            entries = extract_bib_entries(doc, heading_idx)
            for author, journal, year in EXPECTED_ENTRIES:
                if any(check_entry_matches(entry, author, journal, year) for entry in entries):
                    print(f"PASS: Component 2 - Entry for {author} ({journal}, {year}) found (0.10 pts)")
                    total_score += 0.10
                else:
                    print(f"FAIL: Component 2 - Entry for {author} ({journal}, {year}) NOT found in bibliography")
        else:
            print("FAIL: Component 2 - Cannot check entries without bibliography heading")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Entries are sorted alphabetically by author (0.20 points)
    try:
        if heading_idx is not None:
            entries = extract_bib_entries(doc, heading_idx)
            if len(entries) >= 2:
                author_names = []
                for entry in entries:
                    name = get_author_from_entry(entry)
                    if name:
                        author_names.append(name)

                if len(author_names) >= 2:
                    is_sorted = all(author_names[i].lower() <= author_names[i+1].lower()
                                    for i in range(len(author_names) - 1))
                    if is_sorted:
                        print(f"PASS: Component 3 - Entries sorted alphabetically: {author_names} (0.20 pts)")
                        total_score += 0.20
                    else:
                        print(f"FAIL: Component 3 - Entries NOT sorted alphabetically: {author_names}")
                else:
                    print(f"FAIL: Component 3 - Could not extract enough author names from entries")
            else:
                print(f"FAIL: Component 3 - Not enough entries to check sorting ({len(entries)} found)")
        else:
            print("FAIL: Component 3 - Cannot check sorting without bibliography heading")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Exactly 5 bibliography entries (0.15 points)
    try:
        if heading_idx is not None:
            entries = extract_bib_entries(doc, heading_idx)
            if len(entries) == 5:
                print(f"PASS: Component 4 - Exactly 5 bibliography entries found (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 - Expected 5 entries, found {len(entries)}")
        else:
            print("FAIL: Component 4 - Cannot check entry count without bibliography heading")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
