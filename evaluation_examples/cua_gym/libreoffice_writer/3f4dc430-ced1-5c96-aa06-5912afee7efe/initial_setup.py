"""
Initial Setup: Chemistry research paper without bibliography
Task ID: writer_bs_015
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Catalytic Asymmetric Synthesis of Chiral Phosphine Ligands via Palladium-Mediated C\u2013H Activation', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Authors ---
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Dr. Helena Zhao, Prof. Richard Tanaka, Dr. Mei-Lin Wu')
    run.font.size = Pt(11)
    run.font.italic = True

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil.add_run('Department of Chemistry, Stanford University, Stanford, CA 94305, USA')
    run.font.size = Pt(10)

    # --- Abstract ---
    doc.add_heading('Abstract', level=2)
    abstract_text = (
        'We report a novel palladium-catalyzed C\u2013H activation strategy for the enantioselective '
        'synthesis of axially chiral phosphine ligands. Using a newly developed Pd(II)/chiral amino acid '
        'catalyst system, a wide range of biaryl phosphine ligands were obtained in up to 97% ee and 92% '
        'yield. The methodology was demonstrated in the asymmetric Suzuki\u2013Miyaura coupling of '
        'sterically hindered substrates, achieving excellent enantioselectivity (up to 99% ee). Density '
        'functional theory (DFT) calculations provided insights into the stereodetermining transition state, '
        'revealing a concerted metalation\u2013deprotonation (CMD) mechanism.'
    )
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.space_after = Pt(12)

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=2)
    intro_paras = [
        (
            'Chiral phosphine ligands play a central role in transition-metal-catalyzed asymmetric '
            'transformations, enabling precise stereochemical control in numerous bond-forming reactions. '
            'Over the past three decades, the development of new chiral ligand scaffolds has driven '
            'remarkable advances in asymmetric catalysis, with applications spanning pharmaceutical '
            'synthesis, agrochemical production, and materials science.'
        ),
        (
            'Despite these advances, the efficient synthesis of axially chiral biaryl phosphines '
            'remains a significant challenge. Traditional approaches rely on chiral resolution or '
            'stoichiometric chiral auxiliaries, which are inherently wasteful and atom-inefficient. '
            'Recent efforts have focused on catalytic asymmetric methods, particularly C\u2013H '
            'functionalization strategies that enable direct construction of the chiral axis from '
            'readily available starting materials.'
        ),
        (
            'Herein, we present a Pd(II)-catalyzed enantioselective C\u2013H activation approach '
            'for the synthesis of BINAP-type ligands. Our catalyst system, based on a commercially '
            'available Pd(OAc)\u2082 precatalyst and N-acetyl-L-leucine as the chiral ligand, '
            'demonstrates broad substrate scope and excellent functional group tolerance.'
        ),
    ]
    for text in intro_paras:
        doc.add_paragraph(text)

    # --- 2. Results and Discussion ---
    doc.add_heading('2. Results and Discussion', level=2)

    doc.add_heading('2.1 Optimization of Reaction Conditions', level=3)
    opt_text = (
        'We began our investigation by examining the Pd-catalyzed C\u2013H activation of model '
        'substrate 1a (Table 1). Screening of various Pd(II) catalysts revealed that Pd(OAc)\u2082 '
        'in combination with N-acetyl-L-leucine provided the highest enantioselectivity (entry 5, '
        '94% ee). The reaction was sensitive to both temperature and solvent: optimal results were '
        'obtained in 1,2-dichloroethane at 80 \u00b0C over 24 hours. Silver acetate (AgOAc, 2.0 equiv) '
        'served as an effective oxidant.'
    )
    doc.add_paragraph(opt_text)

    # --- Table 1: Optimization results ---
    doc.add_paragraph()  # spacer
    table_title = doc.add_paragraph()
    run = table_title.add_run('Table 1. ')
    run.bold = True
    table_title.add_run('Optimization of reaction conditions for the C\u2013H activation of 1a.')

    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    headers = ['Entry', 'Catalyst (10 mol%)', 'Ligand', 'Yield (%)', 'ee (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['1', 'Pd(OAc)\u2082', 'L-Proline', '45', '32'],
        ['2', 'Pd(OAc)\u2082', 'Boc-L-Val-OH', '62', '58'],
        ['3', 'PdCl\u2082', 'Ac-L-Leu-OH', '71', '82'],
        ['4', 'Pd(OAc)\u2082', 'Ac-D-Leu-OH', '68', '-91'],
        ['5', 'Pd(OAc)\u2082', 'Ac-L-Leu-OH', '87', '94'],
        ['6', 'Pd(TFA)\u2082', 'Ac-L-Leu-OH', '79', '88'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacer

    doc.add_heading('2.2 Substrate Scope', level=3)
    scope_text = (
        'With the optimized conditions in hand, we explored the substrate scope of this '
        'C\u2013H activation protocol (Scheme 2). A variety of aryl groups bearing electron-donating '
        'and electron-withdrawing substituents were tolerated, affording the corresponding chiral '
        'biaryl phosphines in 72\u201392% yield and 88\u201397% ee. Notably, substrates with '
        'ortho-substituents, which are traditionally challenging due to steric effects, reacted '
        'smoothly under our conditions (products 2g\u20132k, 80\u201389% yield, 90\u201395% ee).'
    )
    doc.add_paragraph(scope_text)

    doc.add_heading('2.3 Mechanistic Studies', level=3)
    mech_paras = [
        (
            'To elucidate the mechanism of the enantioselective C\u2013H activation, we performed '
            'a series of kinetic isotope effect (KIE) experiments. A primary KIE of k_H/k_D = 3.8 '
            'was observed, suggesting that C\u2013H bond cleavage is involved in the rate-determining '
            'step. This result is consistent with a concerted metalation\u2013deprotonation (CMD) '
            'pathway.'
        ),
        (
            'DFT calculations at the M06/def2-TZVP level of theory revealed that the stereodetermining '
            'transition state involves a six-membered palladacyclic intermediate. The chiral amino acid '
            'ligand creates a well-defined chiral pocket around the palladium center, directing the '
            'approach of the substrate to one enantiotopic face. The calculated activation barrier '
            'difference between the major and minor diastereomeric transition states (\u0394\u0394G\u2021 = '
            '2.8 kcal/mol) is in good agreement with the experimentally observed enantioselectivity.'
        ),
    ]
    for text in mech_paras:
        doc.add_paragraph(text)

    # --- 3. Applications ---
    doc.add_heading('3. Applications in Asymmetric Catalysis', level=2)
    app_text = (
        'To demonstrate the utility of our chiral phosphine ligands, we applied several '
        'representative products in the Pd-catalyzed asymmetric Suzuki\u2013Miyaura coupling. '
        'Using ligand 2a (3 mol%) with Pd\u2082(dba)\u2083, a range of sterically congested '
        'biaryl products were obtained in 85\u201396% yield and 95\u201399% ee. These results '
        'surpass those obtained with commercially available BINAP under identical conditions '
        '(typically 70\u201385% yield, 80\u201390% ee).'
    )
    doc.add_paragraph(app_text)

    # --- 4. Conclusion ---
    doc.add_heading('4. Conclusion', level=2)
    conclusion_text = (
        'In summary, we have developed a highly enantioselective Pd(II)-catalyzed C\u2013H activation '
        'strategy for the synthesis of axially chiral biaryl phosphine ligands. The method features '
        'mild conditions, broad substrate scope, and excellent enantiocontrol. The synthetic utility '
        'of the products has been demonstrated in asymmetric Suzuki\u2013Miyaura coupling reactions. '
        'Future work will focus on extending this methodology to other classes of chiral ligands and '
        'exploring applications in industrial-scale asymmetric processes.'
    )
    doc.add_paragraph(conclusion_text)

    # --- 5. Experimental Section (abbreviated) ---
    doc.add_heading('5. Experimental Section', level=2)
    exp_text = (
        'General procedure for the enantioselective C\u2013H activation: To a flame-dried Schlenk '
        'tube equipped with a magnetic stir bar were added Pd(OAc)\u2082 (0.05 mmol, 10 mol%), '
        'N-acetyl-L-leucine (0.10 mmol, 20 mol%), AgOAc (1.0 mmol, 2.0 equiv), and substrate '
        '(0.50 mmol, 1.0 equiv) in 1,2-dichloroethane (2.5 mL). The resulting mixture was stirred '
        'at 80 \u00b0C for 24 h under nitrogen atmosphere. After cooling to room temperature, the '
        'reaction mixture was filtered through a pad of Celite and concentrated under reduced '
        'pressure. Purification by column chromatography (silica gel, hexanes/EtOAc gradient) '
        'afforded the desired chiral biaryl phosphine product.'
    )
    doc.add_paragraph(exp_text)

    # No bibliography section -- this is what the task asks the user to add

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
