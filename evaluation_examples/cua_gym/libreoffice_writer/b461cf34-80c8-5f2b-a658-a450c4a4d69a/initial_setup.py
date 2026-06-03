"""
Initial Setup: Play script document with character names and dialogue lines
Task ID: osworld_writer_tabstop_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_006'
OUTPUT = f'{WORKDIR}/play_script_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # 20 lines of dialogue - each line starts with two-word character name
    # No tabstops, no tabs - plain left-aligned text
    lines = [
        "ROMEO MONTAGUE Oh, she doth teach the torches to burn bright",
        "FRIAR LAWRENCE A gentler judgment vanished from his lips",
        "JULIET CAPULET What's in a name? That which we call a rose",
        "BENVOLIO MONTAGUE What sadness lengthens Romeo's hours so long",
        "MERCUTIO MABB Nay, gentle Romeo, we must have you dance",
        "PRINCE ESCALUS Rebellious subjects, enemies to peace",
        "NURSE ANGELICA Come Lammas-eve at night shall she be fourteen",
        "LORD CAPULET But Montague is bound as well as I",
        "LADY CAPULET Verona's summer hath not such a flower",
        "PARIS NOBLE Younger than she are happy mothers made",
        "ROMEO MONTAGUE Did my heart love till now? Forswear it, sight",
        "JULIET CAPULET My only love sprung from my only hate",
        "FRIAR LAWRENCE These violent delights have violent ends",
        "MERCUTIO MABB A plague on both your houses! They have made worms' meat of me",
        "BENVOLIO MONTAGUE I do beseech you, good Mercutio, let's retire",
        "TYBALT CAPULET What, drawn and talk of peace? I hate the word",
        "NURSE ANGELICA Is your man secret? Did you ne'er hear say",
        "LORD CAPULET My child is yet a stranger in the world",
        "FRIAR LAWRENCE Wisely and slow; they stumble that run fast",
        "ROMEO MONTAGUE With love's light wings did I o'er-perch these walls",
    ]

    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(12)
        # No tabstops added - plain left-aligned formatting

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
