"""
Reward Script: Create mailing labels using Avery 5160 label format
Task ID: writer_mt_010
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table/label grid exists in document
  Component 2 (0.25): Table has correct Avery 5160 dimensions (10 rows x 3 cols)
  Component 3 (0.30): All 30 labels contain correct Name/Street/City,State Zip from CSV
  Component 4 (0.20): Page margins adjusted for label layout (not default 1" margins)
"""

import os
import csv

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_010'

def load_csv_records(csv_path):
    """Load contact records from the CSV file."""
    records = []
    try:
        with open(csv_path, 'r') as f:
            reader = csv.DictReader(f)
            for row in reader:
                records.append(row)
    except Exception as e:
        print(f"WARNING: Could not load CSV: {e}")
    return records


def verify_task(file_path, csv_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Load CSV records for content verification
    records = load_csv_records(csv_path)

    # Component 1: Table/label grid exists (0.25 points)
    try:
        num_tables = len(doc.tables)
        if num_tables >= 1:
            print(f"PASS: Component 1 — Document contains {num_tables} table(s) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — No tables found in document (expected label grid)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table has correct Avery 5160 dimensions: 10 rows x 3 columns (0.25 points)
    try:
        if num_tables >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 10 and num_cols == 3:
                print(f"PASS: Component 2 — Table is 10x3 (Avery 5160 layout) (0.25 pts)")
                total_score += 0.25
            elif num_cols == 3 and num_rows >= 10:
                # Partial: right number of columns but extra rows
                print(f"PARTIAL: Component 2 — Table is {num_rows}x{num_cols}, expected 10x3 (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 2 — Table is {num_rows}x{num_cols}, expected 10x3")
        else:
            print(f"FAIL: Component 2 — No table to check dimensions")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Label content matches CSV (0.30 points)
    # Each label should have: Name on line 1, Street on line 2, City, State Zip on line 3
    try:
        if num_tables >= 1 and len(records) > 0:
            table = doc.tables[0]
            matched = 0
            total_labels = len(records)
            label_idx = 0

            for row in table.rows:
                for cell in row.cells:
                    if label_idx >= total_labels:
                        break
                    rec = records[label_idx]
                    cell_text = cell.text.strip()
                    lines = [l.strip() for l in cell_text.split('\n') if l.strip()]

                    if len(lines) >= 3:
                        name_ok = rec['Name'].strip() in lines[0]
                        street_ok = rec['Street'].strip() in lines[1]
                        # City, State Zip format
                        city_state_zip = f"{rec['City']}, {rec['State']} {rec['Zip']}"
                        csz_ok = rec['City'].strip() in lines[2] and rec['State'].strip() in lines[2] and rec['Zip'].strip() in lines[2]

                        if name_ok and street_ok and csz_ok:
                            matched += 1
                        else:
                            print(f"  Label {label_idx} mismatch: expected [{rec['Name']}|{rec['Street']}|{city_state_zip}], got lines={lines[:3]}")
                    else:
                        print(f"  Label {label_idx}: expected 3 lines, got {len(lines)} lines: {lines}")

                    label_idx += 1

            match_ratio = matched / total_labels if total_labels > 0 else 0
            component_score = round(0.30 * match_ratio, 4)
            if matched == total_labels:
                print(f"PASS: Component 3 — All {matched}/{total_labels} labels match CSV data ({component_score} pts)")
            else:
                print(f"PARTIAL: Component 3 — {matched}/{total_labels} labels match CSV data ({component_score} pts)")
            total_score += component_score
        else:
            print(f"FAIL: Component 3 — No table or no CSV records to verify")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Page margins adjusted for Avery 5160 (0.20 points)
    # Avery 5160 requires narrower margins than the default 1" (914400 EMU).
    # The golden uses left/right ~0.1875" (171450 EMU) and top/bottom ~0.5" (457200 EMU).
    # Initial uses default 1" margins. We check that margins are smaller than default.
    try:
        section = doc.sections[0]
        default_margin = 914400  # 1 inch in EMU
        left = section.left_margin
        right = section.right_margin
        top = section.top_margin
        bottom = section.bottom_margin

        # Check that at least left and right margins are reduced from default
        margins_reduced = (left < default_margin) and (right < default_margin)
        if margins_reduced:
            print(f"PASS: Component 4 — Margins adjusted: L={left/914400:.3f}\" R={right/914400:.3f}\" T={top/914400:.3f}\" B={bottom/914400:.3f}\" (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — Margins not reduced from default 1\": L={left/914400:.3f}\" R={right/914400:.3f}\"")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
csv_path = f'{WORKDIR}/Desktop/mailing_list.csv'

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path, csv_path)
