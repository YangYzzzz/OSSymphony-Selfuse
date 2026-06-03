"""
Reward Script: Competitive Analysis Report in LibreOffice Writer
Task ID: writer_wf_057
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15) - Title paragraph present with correct text
  Component 2 (0.20) - Heading 1 sections (6+ required: Exec Summary, Market Overview,
                        Competitor Profiles, Comparative Matrix, Strategic Recommendations, Appendix)
  Component 3 (0.15) - 3 Heading 2 competitor profiles under Competitor Profiles
  Component 4 (0.20) - 3 competitor summary tables with required fields
                        (Company, Market Share, Key Product, Pricing, Strength, Weakness)
  Component 5 (0.20) - Comparative matrix table with 8+ criteria rows and 4 columns
  Component 6 (0.10) - Table of Contents present
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_057'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    tables = doc.tables

    # Gather structural info
    heading1_texts = []
    heading2_texts = []
    title_text = None

    for p in paragraphs:
        style_name = p.style.name if p.style else ''
        if style_name == 'Title':
            title_text = p.text.strip()
        elif style_name == 'Heading 1':
            heading1_texts.append(p.text.strip())
        elif style_name == 'Heading 2':
            heading2_texts.append(p.text.strip())

    print(f"INFO: Found title='{title_text}', {len(heading1_texts)} H1, {len(heading2_texts)} H2, {len(tables)} tables")
    print(f"INFO: H1 sections: {heading1_texts}")
    print(f"INFO: H2 sections: {heading2_texts}")

    # Component 1: Title paragraph (0.15 points)
    try:
        if title_text and 'competitive' in title_text.lower() and 'cloud storage' in title_text.lower():
            print(f"PASS: Component 1 — Title found: '{title_text}' (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Expected title with 'Competitive' and 'Cloud Storage', found: '{title_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Heading 1 sections (0.20 points)
    # Required: Executive Summary, Market Overview, Competitor Profiles,
    #           Comparative Matrix, Strategic Recommendations, Appendix
    try:
        required_h1_keywords = [
            'executive summary',
            'market overview',
            'competitor profile',
            'comparative matrix',
            'strategic recommendation',
            'appendix',
        ]
        h1_lower = [h.lower() for h in heading1_texts]
        matched_h1 = 0
        for kw in required_h1_keywords:
            if any(kw in h for h in h1_lower):
                matched_h1 += 1

        if matched_h1 >= 6:
            print(f"PASS: Component 2 — All 6 required H1 sections found (0.20 pts)")
            total_score += 0.20
        elif matched_h1 >= 4:
            partial = round(0.20 * (matched_h1 / 6), 2)
            print(f"PARTIAL: Component 2 — {matched_h1}/6 required H1 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched_h1}/6 required H1 sections found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 3 Heading 2 competitor profiles (0.15 points)
    try:
        # Need at least 3 Heading 2 entries that are competitor names
        if len(heading2_texts) >= 3:
            print(f"PASS: Component 3 — {len(heading2_texts)} H2 competitor profiles found (0.15 pts)")
            total_score += 0.15
        elif len(heading2_texts) >= 1:
            partial = round(0.15 * (len(heading2_texts) / 3), 2)
            print(f"PARTIAL: Component 3 — {len(heading2_texts)}/3 H2 competitor profiles ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No H2 competitor profiles found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 3 competitor summary tables with required fields (0.20 points)
    try:
        required_fields = {'company', 'market share', 'key product', 'pricing', 'strength', 'weakness'}
        valid_summary_tables = 0

        for ti, table in enumerate(tables):
            # A summary table has ~6 rows x 2 cols with label-value pairs
            if len(table.columns) == 2 and len(table.rows) >= 5:
                field_labels = set()
                for row in table.rows:
                    label = row.cells[0].text.strip().lower()
                    field_labels.add(label)
                # Check how many required fields are present
                matches = required_fields.intersection(field_labels)
                if len(matches) >= 4:
                    valid_summary_tables += 1
                    print(f"  INFO: Table {ti} is a valid summary table ({len(matches)}/6 required fields)")

        if valid_summary_tables >= 3:
            print(f"PASS: Component 4 — {valid_summary_tables} valid competitor summary tables (0.20 pts)")
            total_score += 0.20
        elif valid_summary_tables >= 1:
            partial = round(0.20 * (valid_summary_tables / 3), 2)
            print(f"PARTIAL: Component 4 — {valid_summary_tables}/3 valid summary tables ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No valid competitor summary tables found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Comparative matrix table (0.20 points)
    # Should have 4 columns (Criteria + 3 competitors) and 8+ data rows (+ header = 9+ rows)
    try:
        matrix_found = False
        for ti, table in enumerate(tables):
            if len(table.columns) >= 4 and len(table.rows) >= 9:
                # Check header row has "Criteria" or similar
                header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if any('criteri' in c for c in header_cells):
                    # Count data rows (non-empty)
                    data_rows = sum(1 for r in table.rows[1:] if r.cells[0].text.strip())
                    if data_rows >= 8:
                        print(f"PASS: Component 5 — Comparative matrix found (table {ti}, {data_rows} criteria rows, {len(table.columns)} cols) (0.20 pts)")
                        total_score += 0.20
                        matrix_found = True
                        break
                    elif data_rows >= 5:
                        partial = round(0.20 * (data_rows / 8), 2)
                        print(f"PARTIAL: Component 5 — Matrix found but only {data_rows}/8 criteria rows ({partial} pts)")
                        total_score += partial
                        matrix_found = True
                        break

        if not matrix_found:
            # Fallback: look for any table with 4+ cols and multiple rows
            for ti, table in enumerate(tables):
                if len(table.columns) >= 4 and len(table.rows) >= 5:
                    data_rows = sum(1 for r in table.rows[1:] if r.cells[0].text.strip())
                    if data_rows >= 4:
                        partial = round(0.20 * 0.5, 2)
                        print(f"PARTIAL: Component 5 — Possible matrix table {ti} ({data_rows} rows, {len(table.columns)} cols) ({partial} pts)")
                        total_score += partial
                        matrix_found = True
                        break

            if not matrix_found:
                print(f"FAIL: Component 5 — No comparative matrix table found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Table of Contents (0.10 points)
    # TOC can be a Heading 1 "Table of Contents" or similar marker
    try:
        toc_found = False
        for p in paragraphs:
            text_lower = p.text.strip().lower()
            if 'table of contents' in text_lower or 'toc' == text_lower:
                toc_found = True
                break
            # Also check for TOC field codes in XML
            if p._element.xml and 'w:fldChar' in p._element.xml and 'TOC' in p._element.xml:
                toc_found = True
                break

        if toc_found:
            print(f"PASS: Component 6 — Table of Contents found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — No Table of Contents found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
