"""
Reward Script: Remove all duplicate entries from the glossary list in a Writer document.
Task ID: osworld_writer_duplicate_line_removal_003
Domain: libreoffice_writer
Scoring:
  Component 1 — Glossary paragraph count == 8 (0.3 pts)
  Component 2 — No duplicate terms in glossary section (0.4 pts)
  Component 3 — Terms appear in correct first-occurrence order (0.3 pts)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_duplicate_line_removal_003'

# Expected 8 unique glossary terms in their original first-occurrence order
EXPECTED_TERMS_ORDER = [
    'API',
    'Agile',
    'CI/CD',
    'Deployment',
    'Encapsulation',
    'Framework',
    'Git',
    'HTTP',
]


def get_glossary_paragraphs(doc):
    """
    Extract glossary entry paragraphs. The document structure is:
    [0] Title, [1] Normal (description), [2] Heading 1 ('Glossary Terms'),
    [3..] Normal paragraphs that are glossary entries.
    Returns list of text strings for all Normal paragraphs after Heading 1.
    """
    heading_found = any(
        p.style.name == 'Heading 1' and 'Glossary' in p.text
        for p in doc.paragraphs
    )
    if not heading_found:
        return []

    # Find the index of the Glossary heading and collect entries after it
    entries = []
    past_heading = False
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1' and 'Glossary' in para.text:
            past_heading = (para.text.strip() != '')  # set True when we encounter heading
        elif past_heading and para.style.name == 'Normal' and para.text.strip():
            entries.append(para.text.strip())
    return entries


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract glossary entries (paragraphs in the glossary section)
    try:
        glossary_entries = get_glossary_paragraphs(doc)
        print(f"INFO: Found {len(glossary_entries)} glossary entry paragraphs")
        for i, entry in enumerate(glossary_entries):
            print(f"  Entry [{i}]: {entry[:60]!r}")
    except Exception as e:
        print(f"ERROR: Could not extract glossary entries: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Glossary section has exactly 8 entries (0.3 points)
    # Initial has 12 glossary entries; golden should have 8 after removing 4 duplicates.
    try:
        expected_count = 8
        actual_count = len(glossary_entries)
        if actual_count == expected_count:
            print(f"PASS: Component 1 — Glossary has {actual_count} entries (expected {expected_count}) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Glossary has {actual_count} entries, expected {expected_count}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: No duplicate terms in the glossary section (0.4 points)
    # Each term (identified by the prefix before ':') should appear exactly once.
    try:
        def extract_term(entry):
            """Extract the key term before the colon."""
            if ':' in entry:
                return entry.split(':')[0].strip()
            return entry.strip()

        terms = [extract_term(e) for e in glossary_entries]
        seen = {}
        duplicates = []
        for t in terms:
            if t in seen:
                duplicates.append(t)
            else:
                seen[t] = True

        if len(duplicates) == 0:
            print(f"PASS: Component 2 — No duplicate terms found in glossary (0.4 pts)")
            print(f"  Terms: {terms}")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Found duplicate terms: {duplicates}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Terms appear in correct first-occurrence order (0.3 points)
    # The 8 unique terms should appear in the order of their first occurrence in the original:
    # API, Agile, CI/CD, Deployment, Encapsulation, Framework, Git, HTTP
    try:
        def extract_term(entry):
            if ':' in entry:
                return entry.split(':')[0].strip()
            return entry.strip()

        actual_terms = [extract_term(e) for e in glossary_entries]
        if actual_terms == EXPECTED_TERMS_ORDER:
            print(f"PASS: Component 3 — Terms in correct first-occurrence order (0.3 pts)")
            print(f"  Order: {actual_terms}")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Expected order {EXPECTED_TERMS_ORDER}, found {actual_terms}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
