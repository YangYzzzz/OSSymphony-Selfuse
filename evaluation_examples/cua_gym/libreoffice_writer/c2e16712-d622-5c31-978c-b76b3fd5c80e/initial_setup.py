"""
Initial Setup: Glossary document with duplicate entries for duplicate removal task
Task ID: osworld_writer_duplicate_line_removal_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_duplicate_line_removal_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Software Development Glossary', level=0)

    # --- Introduction paragraph ---
    intro = doc.add_paragraph(
        'This glossary provides definitions for common terms used in software development, '
        'project management, and systems architecture. Refer to this document for consistent '
        'terminology across all project documentation and communications.'
    )

    # --- Glossary heading ---
    doc.add_heading('Glossary Terms', level=1)

    # --- 12 glossary lines (8 unique, 4 duplicated) ---
    # Order: 12 total lines with 4 duplicates scattered throughout
    # Unique terms: API, Agile, CI/CD, Deployment, Encapsulation, Framework, Git, HTTP
    # Duplicated terms (appear twice): API, Agile, CI/CD, Deployment
    glossary_lines = [
        'API: An Application Programming Interface that defines interactions between software components.',
        'Agile: An iterative software development methodology emphasizing flexibility and customer collaboration.',
        'CI/CD: Continuous Integration and Continuous Delivery, a practice of automating build and deployment pipelines.',
        'Deployment: The process of releasing and distributing a software application to a target environment.',
        'Encapsulation: An object-oriented principle that bundles data and methods within a single unit.',
        'API: An Application Programming Interface that defines interactions between software components.',
        'Framework: A reusable set of libraries and tools that provide a foundation for building applications.',
        'Agile: An iterative software development methodology emphasizing flexibility and customer collaboration.',
        'Git: A distributed version control system for tracking changes in source code during development.',
        'CI/CD: Continuous Integration and Continuous Delivery, a practice of automating build and deployment pipelines.',
        'HTTP: HyperText Transfer Protocol, the foundation of data communication on the World Wide Web.',
        'Deployment: The process of releasing and distributing a software application to a target environment.',
    ]

    for line in glossary_lines:
        para = doc.add_paragraph(line)
        para.style = doc.styles['Normal']

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Glossary lines: 12 total (8 unique, 4 duplicated)')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
