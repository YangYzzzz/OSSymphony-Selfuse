"""
Initial Setup: Document with 20 pages, all using Arabic page numbering.
Task ID: writer_fs_064
Domain: libreoffice_writer

Creates a Writer document with:
- Page 1: Title page (front matter)
- Page 2: Table of Contents (front matter)
- Pages 3-20: Main body content (18 pages of realistic report text)
- All pages numbered with Arabic numerals 1-20 in the footer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_064'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph for page numbering."""
    run = paragraph.add_run()
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fldChar2)


def setup_footer_with_page_number(section):
    """Add centered page number to footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear existing paragraphs
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp)


def add_body_paragraphs(doc, heading, paragraphs_text):
    """Add a heading and multiple paragraphs of body text."""
    doc.add_heading(heading, level=2)
    for text in paragraphs_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)


def create_initial():
    doc = Document()

    # ---- Page Setup ----
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ---- PAGE 1: Title Page ----
    # Add some vertical spacing
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    title = doc.add_heading('Annual Performance Review Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Fiscal Year 2024-2025')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    org = doc.add_paragraph()
    org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = org.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(14)

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run('Human Resources Division')
    run.font.size = Pt(12)
    run.font.italic = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_p.paragraph_format.space_before = Pt(36)
    run = date_p.add_run('Prepared: March 15, 2025')
    run.font.size = Pt(11)

    conf = doc.add_paragraph()
    conf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = conf.add_run('CONFIDENTIAL')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Page break after title page
    doc.add_page_break()

    # ---- PAGE 2: Table of Contents ----
    toc_heading = doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ('1. Executive Summary', '3'),
        ('2. Methodology', '5'),
        ('3. Department Performance Overview', '7'),
        ('   3.1 Engineering Division', '7'),
        ('   3.2 Marketing & Sales', '9'),
        ('   3.3 Operations & Logistics', '10'),
        ('4. Individual Performance Metrics', '11'),
        ('5. Training & Development Initiatives', '13'),
        ('6. Compensation Analysis', '15'),
        ('7. Retention Statistics', '16'),
        ('8. Recommendations', '17'),
        ('9. Appendix A: Raw Survey Data', '19'),
        ('10. Appendix B: Statistical Methods', '20'),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        if not item.startswith('   '):
            run.font.size = Pt(11)
        else:
            run.font.size = Pt(10)
        # Add dots and page number
        run2 = p.add_run(f'  {"." * 40}  {page}')
        run2.font.size = Pt(10)

    # Page break after TOC
    doc.add_page_break()

    # ---- PAGES 3-20: Main Body (18 pages of content) ----
    # We need to generate enough content to fill 18 pages
    body_sections = [
        ('1. Executive Summary', [
            'This annual performance review report provides a comprehensive analysis of employee performance across all divisions of Meridian Technologies Inc. for the fiscal year 2024-2025. The evaluation period covers April 1, 2024 through March 31, 2025.',
            'Key findings indicate an overall improvement of 12.3% in productivity metrics compared to the previous fiscal year. The Engineering Division demonstrated exceptional growth, with project completion rates increasing by 18.7%. Customer satisfaction scores across the organization averaged 4.2 out of 5.0, representing a significant improvement from last year\'s 3.8 average.',
            'Employee engagement survey results show that 78% of staff report high satisfaction with their work environment, while 82% feel adequately supported in their professional development goals. The voluntary turnover rate decreased from 14.2% to 11.8%, saving an estimated $2.3 million in recruitment and training costs.',
            'Total headcount grew from 847 to 923 employees during the review period, with the most significant expansion occurring in the Engineering and Data Science teams. The organization successfully onboarded 142 new hires while maintaining quality standards.',
            'This report details performance metrics by department, highlights individual achievements, analyzes compensation trends, and provides strategic recommendations for the upcoming fiscal year. All data has been verified by the Internal Audit team.',
        ]),
        ('2. Methodology', [
            'The performance evaluation framework employed this year builds upon our established Balanced Scorecard methodology, incorporating 360-degree feedback mechanisms and objective key results (OKR) tracking. Each employee was assessed across four primary dimensions: Technical Competency, Collaboration & Leadership, Innovation & Initiative, and Client Impact.',
            'Data collection methods included quarterly manager assessments, peer review surveys distributed through our internal HR platform, client feedback aggregation from our CRM system, and automated productivity tracking through project management tools. All quantitative data was normalized using standard deviation scaling to ensure fair cross-departmental comparison.',
            'The statistical analysis was performed using a combination of descriptive statistics for trend identification and inferential methods including ANOVA for inter-departmental comparisons and regression analysis for predictive modeling of performance trajectories.',
            'Survey instruments were validated through pilot testing with a sample group of 45 employees across departments. Cronbach\'s alpha for internal consistency measured 0.87, exceeding the acceptable threshold of 0.70. Response rates averaged 91% across all quarterly surveys.',
            'Performance ratings use a five-point scale: Exceptional (5), Exceeds Expectations (4), Meets Expectations (3), Needs Improvement (2), and Unsatisfactory (1). Calibration sessions were conducted with all department heads to ensure rating consistency and minimize bias.',
        ]),
        ('3. Department Performance Overview', [
            'This section provides a detailed breakdown of performance metrics across the three major divisions of Meridian Technologies Inc. Each department was evaluated on division-specific KPIs in addition to the organization-wide metrics described in the methodology section.',
        ]),
        ('3.1 Engineering Division', [
            'The Engineering Division, led by VP of Engineering Dr. Priya Sharma, demonstrated outstanding performance during the review period. The division comprises 312 employees across six teams: Frontend Development, Backend Infrastructure, Mobile Applications, DevOps & Cloud, Quality Assurance, and Data Engineering.',
            'Sprint velocity improved by 23% year-over-year, with the average story points completed per two-week sprint increasing from 42 to 52. Code review turnaround time decreased from an average of 18 hours to 11 hours, indicating improved collaboration and process efficiency.',
            'The team successfully delivered 14 major product releases during the fiscal year, including the highly anticipated Platform 3.0 migration which was completed two weeks ahead of schedule. Critical bug density in production decreased by 31%, from 2.3 to 1.6 defects per thousand lines of code.',
            'Notable achievements include the successful implementation of the microservices architecture migration, reducing average API response time from 340ms to 89ms. The cloud infrastructure optimization project led by Senior Engineer James Rodriguez resulted in a 28% reduction in AWS costs, saving approximately $1.2 million annually.',
            'Areas for improvement include cross-team documentation practices and knowledge transfer processes. Exit interviews from the Engineering Division cited insufficient documentation as a contributing factor in 3 of 8 voluntary departures.',
        ]),
        ('3.2 Marketing & Sales', [
            'The Marketing & Sales Division, under the leadership of Chief Revenue Officer Michael Torres, comprises 198 employees across Sales Operations, Digital Marketing, Brand Strategy, Customer Success, and Business Development teams.',
            'Revenue targets were exceeded by 8.4%, with total annual revenue reaching $127.3 million against a target of $117.4 million. The average deal size increased from $34,200 to $41,800, reflecting the sales team\'s successful shift toward enterprise-level accounts. Customer acquisition cost decreased by 15% through improved digital marketing targeting.',
            'The Digital Marketing team achieved a 45% increase in qualified leads generated through organic channels, reducing dependency on paid advertising. Content marketing efforts produced 156 blog posts, 24 whitepapers, and 12 webinars, with an average engagement rate of 4.7%.',
            'Customer retention rate improved from 88.2% to 91.7%, with the Customer Success team implementing a proactive outreach program that identified at-risk accounts 30 days earlier than the previous system. Net Promoter Score (NPS) increased from 42 to 58.',
        ]),
        ('3.3 Operations & Logistics', [
            'The Operations & Logistics Division, managed by COO Sandra Williams, includes 413 employees across Supply Chain Management, Warehouse Operations, Transportation, Procurement, and Facilities Management.',
            'Operational efficiency metrics showed consistent improvement throughout the year. Order fulfillment accuracy reached 99.4%, up from 98.1%. Average order processing time decreased from 2.3 business days to 1.7 business days. Inventory turnover ratio improved from 8.2 to 9.6.',
            'The division successfully implemented a new warehouse management system (WMS) that reduced picking errors by 62% and improved inventory tracking accuracy to 99.8%. The implementation was completed on budget and within the projected six-month timeline.',
            'Sustainability initiatives led by the Operations team resulted in a 22% reduction in carbon emissions from logistics operations. The transition to electric delivery vehicles in metropolitan areas saved $340,000 in fuel costs while reducing the fleet\'s carbon footprint by 180 metric tons.',
        ]),
        ('4. Individual Performance Metrics', [
            'This section highlights individual performance data aggregated across the organization. During the review period, 923 employees received formal performance evaluations. The distribution of ratings was as follows: Exceptional - 12% (111 employees), Exceeds Expectations - 34% (314 employees), Meets Expectations - 41% (378 employees), Needs Improvement - 11% (101 employees), Unsatisfactory - 2% (19 employees).',
            'The top performers across divisions included Sarah Chen (Engineering) who led the Platform 3.0 migration, achieving 147% of her OKR targets. Marcus Johnson (Marketing) generated $18.7 million in new business, the highest individual contribution in company history. Elena Vasquez (Operations) redesigned the warehouse layout, improving throughput by 34%.',
            'Performance improvement plans (PIPs) were initiated for 38 employees, with a 63% success rate at the six-month review point. The most common areas requiring improvement were communication skills (42%), deadline management (28%), and technical skill gaps (30%).',
            'Promotion rates remained stable at 14.6%, with 135 employees advancing to higher positions. Internal mobility accounted for 23% of all role changes, with employees moving between departments to align with career development goals and organizational needs.',
        ]),
        ('5. Training & Development Initiatives', [
            'The Learning & Development department invested $2.8 million in employee training programs during the fiscal year, representing a 20% increase from the previous year. A total of 4,230 training hours were completed across the organization, averaging 4.6 hours per employee per month.',
            'Key programs included the Technical Leadership Academy (87 participants), the Sales Excellence Bootcamp (62 participants), and the Operations Management Certificate Program (45 participants). Post-training assessment scores showed an average knowledge retention improvement of 34%.',
            'The new mentorship program paired 78 junior employees with senior leaders for six-month mentoring cycles. Participant satisfaction averaged 4.5 out of 5.0, with 89% of mentees reporting significant career clarity improvements. The program will be expanded by 40% in the upcoming fiscal year.',
            'External certification support resulted in 156 employees obtaining industry-recognized certifications, including AWS Solutions Architect (42), PMP (28), Six Sigma Green Belt (31), and various vendor-specific technical certifications (55).',
        ]),
        ('6. Compensation Analysis', [
            'The total compensation budget for the review period was $89.4 million, representing 70.3% of revenue. Base salary adjustments averaged 4.2% across the organization, exceeding the national average of 3.6%. Merit-based bonuses totaled $8.7 million, with an average payout of $9,425 per eligible employee.',
            'Market competitiveness analysis conducted in partnership with Willis Towers Watson indicated that Meridian\'s total compensation packages rank in the 68th percentile for the technology sector. Engineering roles were positioned at the 75th percentile, while operations roles were at the 55th percentile.',
            'Equity compensation through the employee stock option program represented 12% of total compensation for senior staff. The company granted 450,000 new stock options during the period, with a weighted average exercise price of $42.30.',
            'Benefits utilization data showed that 94% of employees enrolled in the health insurance plan, 78% participated in the 401(k) program with an average contribution rate of 7.2%, and 62% utilized the flexible spending account benefit.',
        ]),
        ('7. Retention Statistics', [
            'Overall employee retention rate for the fiscal year was 88.2%, an improvement of 2.4 percentage points from the previous year. Voluntary turnover decreased from 14.2% to 11.8%, while involuntary turnover remained stable at approximately 2.0%.',
            'Exit interview analysis revealed the top three reasons for voluntary departure: career advancement opportunities at other organizations (38%), compensation competitiveness (27%), and work-life balance concerns (19%). Geographic relocation accounted for 11% of departures.',
            'Retention was strongest in the Engineering Division (91.3%) and weakest in the Customer Success team (83.4%). Targeted retention initiatives for the Customer Success team include revised career pathing, additional training budgets, and flexible scheduling options.',
            'The average tenure at the company increased from 3.4 years to 3.7 years. Employees who participated in training programs showed 23% lower turnover rates than non-participants.',
        ]),
        ('8. Recommendations', [
            'Based on the comprehensive analysis presented in this report, the following strategic recommendations are proposed for the upcoming fiscal year:',
            'First, increase investment in technical documentation and knowledge management systems. The Engineering Division should allocate 10% of sprint capacity to documentation tasks, and a dedicated Technical Writer should be hired to support cross-team knowledge sharing.',
            'Second, expand the mentorship program to include reverse mentoring, where junior employees share emerging technology knowledge with senior leaders. This bidirectional approach has shown 28% higher satisfaction rates in peer organizations.',
            'Third, address the compensation gap in Operations & Logistics roles by conducting a targeted market adjustment. A 6-8% salary increase for key operational roles would bring compensation to the 65th percentile and is projected to reduce turnover by 3-4 percentage points.',
            'Fourth, implement a quarterly pulse survey to complement the annual engagement assessment. Real-time feedback mechanisms enable faster identification and resolution of emerging workplace issues.',
            'Fifth, establish a formal succession planning process for all director-level and above positions. Currently, only 40% of leadership roles have identified successors, posing organizational risk.',
        ]),
        ('9. Appendix A: Raw Survey Data', [
            'Employee Engagement Survey Results Summary (Q1-Q4 2024-2025):',
            'Q1 Results - Overall Satisfaction: 3.9/5.0, Work Environment: 4.1/5.0, Management Support: 3.8/5.0, Career Growth: 3.6/5.0, Compensation Fairness: 3.5/5.0. Response Rate: 89%. Total Respondents: 753.',
            'Q2 Results - Overall Satisfaction: 4.0/5.0, Work Environment: 4.2/5.0, Management Support: 3.9/5.0, Career Growth: 3.7/5.0, Compensation Fairness: 3.6/5.0. Response Rate: 92%. Total Respondents: 798.',
            'Q3 Results - Overall Satisfaction: 4.1/5.0, Work Environment: 4.2/5.0, Management Support: 4.0/5.0, Career Growth: 3.8/5.0, Compensation Fairness: 3.7/5.0. Response Rate: 90%. Total Respondents: 812.',
            'Q4 Results - Overall Satisfaction: 4.2/5.0, Work Environment: 4.3/5.0, Management Support: 4.1/5.0, Career Growth: 3.9/5.0, Compensation Fairness: 3.8/5.0. Response Rate: 93%. Total Respondents: 845.',
        ]),
        ('10. Appendix B: Statistical Methods', [
            'Performance Score Normalization: All raw performance scores were converted to z-scores using the formula z = (x - mu) / sigma, where mu represents the departmental mean and sigma represents the departmental standard deviation. This normalization ensures equitable comparison across departments with different scoring tendencies.',
            'ANOVA Analysis: One-way ANOVA was used to test for significant differences in performance ratings across departments. The F-statistic of 3.47 (p = 0.012) indicated significant inter-departmental variation, leading to post-hoc Tukey HSD tests for pairwise comparisons.',
            'Regression Modeling: A multiple linear regression model was developed to predict performance ratings based on tenure, training hours, prior year rating, and department. The model achieved an adjusted R-squared of 0.64, with training hours (beta = 0.31, p < 0.001) and prior year rating (beta = 0.42, p < 0.001) as the strongest predictors.',
            'Survey Reliability: Internal consistency was assessed using Cronbach\'s alpha for each survey dimension. All dimensions exceeded the 0.70 threshold: Overall Satisfaction (0.89), Work Environment (0.85), Management Support (0.87), Career Growth (0.82), Compensation Fairness (0.84).',
        ]),
    ]

    for i, (heading, paragraphs) in enumerate(body_sections):
        if i > 0:
            # Add spacing between sections but not page breaks for every section
            # Let content flow naturally
            pass
        add_body_paragraphs(doc, heading, paragraphs)
        # Add page breaks after certain sections to ensure we get ~18 pages
        if heading in ['1. Executive Summary', '3.1 Engineering Division',
                       '3.3 Operations & Logistics', '5. Training & Development Initiatives',
                       '7. Retention Statistics', '8. Recommendations']:
            doc.add_page_break()

    # ---- Setup footer with Arabic page numbers for ALL pages ----
    # Only one section, so all pages get Arabic numbers 1-20
    setup_footer_with_page_number(doc.sections[0])

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
