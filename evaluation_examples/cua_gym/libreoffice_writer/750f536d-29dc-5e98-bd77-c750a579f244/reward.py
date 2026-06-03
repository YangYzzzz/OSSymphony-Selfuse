"""
Reward Script: Set line spacing to 1.5 lines for entire document
Task ID: writer_tech_009
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Normal paragraphs have 1.5 line spacing
  Component 2 (0.3): Heading paragraphs have 1.5 line spacing
  Component 3 (0.3): List Bullet paragraphs have 1.5 line spacing
"""

import os

from docx import Document
from docx.enum.text import WD_LINE_SPACING

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_009'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def check_line_spacing_15(para):
    """
    Check if a paragraph has 1.5 line spacing.
    Accepts either:
      - line_spacing == 1.5 with rule PROPORTIONAL or MULTIPLE
      - line_spacing_rule == ONE_POINT_FIVE (enum value 1)
    """
    pf = para.paragraph_format
    ls = pf.line_spacing
    rule = pf.line_spacing_rule

    # Direct rule check
    if rule is not None and rule == WD_LINE_SPACING.ONE_POINT_FIVE:
        return True

    # Proportional/multiple 1.5
    if ls is not None and abs(float(ls) - 1.5) < 0.01:
        return True

    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Classify paragraphs by style
    normal_paras = []
    heading_paras = []
    bullet_paras = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'
        if style_name.startswith('Heading') or style_name == 'Title':
            heading_paras.append(para)
        elif style_name.startswith('List'):
            bullet_paras.append(para)
        else:
            normal_paras.append(para)

    print(f"Document has {len(doc.paragraphs)} paragraphs: "
          f"{len(normal_paras)} Normal, {len(heading_paras)} Heading/Title, "
          f"{len(bullet_paras)} List")

    # Component 1: Normal paragraphs have 1.5 line spacing (0.4 points)
    try:
        if len(normal_paras) == 0:
            print("FAIL: Component 1 -- No Normal paragraphs found")
        else:
            passed = sum(1 for p in normal_paras if check_line_spacing_15(p))
            ratio = passed / len(normal_paras)
            if ratio >= 0.9:
                print(f"PASS: Component 1 -- {passed}/{len(normal_paras)} Normal paragraphs have 1.5 spacing (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 -- Only {passed}/{len(normal_paras)} Normal paragraphs have 1.5 spacing")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Heading/Title paragraphs have 1.5 line spacing (0.3 points)
    try:
        if len(heading_paras) == 0:
            print("FAIL: Component 2 -- No Heading paragraphs found")
        else:
            passed = sum(1 for p in heading_paras if check_line_spacing_15(p))
            ratio = passed / len(heading_paras)
            if ratio >= 0.9:
                print(f"PASS: Component 2 -- {passed}/{len(heading_paras)} Heading/Title paragraphs have 1.5 spacing (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 -- Only {passed}/{len(heading_paras)} Heading/Title paragraphs have 1.5 spacing")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: List Bullet paragraphs have 1.5 line spacing (0.3 points)
    try:
        if len(bullet_paras) == 0:
            print("FAIL: Component 3 -- No List Bullet paragraphs found")
        else:
            passed = sum(1 for p in bullet_paras if check_line_spacing_15(p))
            ratio = passed / len(bullet_paras)
            if ratio >= 0.9:
                print(f"PASS: Component 3 -- {passed}/{len(bullet_paras)} List Bullet paragraphs have 1.5 spacing (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 -- Only {passed}/{len(bullet_paras)} List Bullet paragraphs have 1.5 spacing")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification (task may have unsaved GUI edits)
persist_app_state("libreoffice_writer")

# Run verification
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
