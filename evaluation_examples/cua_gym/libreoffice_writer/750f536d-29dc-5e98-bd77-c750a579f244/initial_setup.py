"""
Initial Setup: Technical documentation with single line spacing
Task ID: writer_tech_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style to single spacing
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(6)

    # --- Title ---
    title = doc.add_heading('CloudSync API Integration Guide', level=0)
    title.paragraph_format.line_spacing = 1.0

    # --- Section 1 ---
    h1 = doc.add_heading('1. Overview', level=1)
    h1.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'CloudSync is a RESTful API service designed for real-time data synchronization '
        'across distributed cloud environments. This guide covers the integration workflow, '
        'authentication mechanisms, and best practices for production deployments. The API '
        'supports JSON and Protocol Buffers serialization formats, with automatic content '
        'negotiation based on the Accept header.'
    )
    p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'The service handles approximately 2.4 million requests per day across three '
        'geographic regions: US-East (Virginia), EU-West (Frankfurt), and AP-Southeast '
        '(Singapore). Average response latency is 47ms at the 95th percentile, with '
        'automatic failover between regions when health checks detect degradation.'
    )
    p.paragraph_format.line_spacing = 1.0

    # --- Section 2 ---
    h2 = doc.add_heading('2. Authentication', level=1)
    h2.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'All API requests require OAuth 2.0 bearer tokens issued by the CloudSync Identity '
        'Provider. Tokens expire after 3600 seconds and must be refreshed using the refresh '
        'token grant flow. Client credentials are provisioned through the developer portal '
        'at https://developer.cloudsync.io/credentials.'
    )
    p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'For service-to-service communication, mutual TLS (mTLS) authentication is '
        'recommended. Certificate rotation is handled automatically by the CloudSync '
        'certificate manager, which issues new certificates 72 hours before expiration. '
        'The root CA certificate chain must be installed in the client trust store.'
    )
    p.paragraph_format.line_spacing = 1.0

    # --- Section 3 ---
    h3 = doc.add_heading('3. Core Endpoints', level=1)
    h3.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'The synchronization engine exposes four primary endpoints for data management. '
        'Each endpoint accepts standard HTTP methods and returns responses in the negotiated '
        'format. Rate limiting is enforced at 1000 requests per minute per API key, with '
        'burst allowance up to 150 requests in a 10-second window.'
    )
    p.paragraph_format.line_spacing = 1.0

    # Table of endpoints
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Endpoint', 'Method', 'Description']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    endpoints = [
        ['/api/v2/sync', 'POST', 'Initiate a synchronization job between source and target'],
        ['/api/v2/status/{job_id}', 'GET', 'Retrieve current status and progress of a sync job'],
        ['/api/v2/conflicts', 'GET', 'List unresolved data conflicts requiring manual resolution'],
        ['/api/v2/webhooks', 'PUT', 'Register or update webhook callbacks for job events'],
    ]
    for r, row_data in enumerate(endpoints, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Section 4 ---
    h4 = doc.add_heading('4. Error Handling', level=1)
    h4.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'The API uses standard HTTP status codes for error signaling. Responses with status '
        'codes 4xx indicate client-side issues such as malformed requests, expired tokens, '
        'or insufficient permissions. Status codes 5xx indicate server-side failures that '
        'should trigger automatic retry with exponential backoff.'
    )
    p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'Retry logic should implement jittered exponential backoff starting at 500ms with a '
        'maximum delay of 32 seconds. The Retry-After header, when present, takes precedence '
        'over calculated backoff intervals. Circuit breaker patterns are recommended for '
        'production integrations to prevent cascading failures during regional outages.'
    )
    p.paragraph_format.line_spacing = 1.0

    # --- Section 5 ---
    h5 = doc.add_heading('5. Performance Optimization', level=1)
    h5.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'To minimize latency, clients should maintain persistent HTTP/2 connections and use '
        'connection pooling with a minimum pool size of 4 connections per region. Request '
        'batching is supported via the /api/v2/sync/batch endpoint, which accepts up to 50 '
        'synchronization requests in a single HTTP call, reducing overhead significantly.'
    )
    p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph(
        'Data compression using gzip or Brotli encoding reduces payload size by approximately '
        '78% for typical JSON responses. Enable compression by setting the Accept-Encoding '
        'header. The API server automatically compresses responses larger than 1024 bytes. '
        'For Protocol Buffers payloads, additional compression yields diminishing returns '
        'due to the already compact binary format.'
    )
    p.paragraph_format.line_spacing = 1.0

    # --- Section 6 ---
    h6 = doc.add_heading('6. Deployment Checklist', level=1)
    h6.paragraph_format.line_spacing = 1.0

    items = [
        'Verify OAuth 2.0 credentials are provisioned and stored securely in a vault',
        'Configure mTLS certificates for service-to-service endpoints',
        'Set up monitoring dashboards for sync job success rates and latency',
        'Implement retry logic with jittered exponential backoff',
        'Register webhook endpoints for real-time job completion notifications',
        'Load test the integration at 120% of projected peak traffic volume',
        'Configure alerting thresholds for error rate exceeding 0.5%',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing = 1.0

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
