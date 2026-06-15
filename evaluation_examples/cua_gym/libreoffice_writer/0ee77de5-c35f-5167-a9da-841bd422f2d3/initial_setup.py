"""
Initial Setup: Find all Fahrenheit temperatures and append Celsius equivalents
Task ID: writer_frd_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Cookbook Recipes', level=0)

    # Introduction
    intro = doc.add_paragraph(
        'Welcome to our family cookbook! This collection features beloved recipes '
        'passed down through generations. Each dish has been tested and perfected '
        'over the years. Please follow the temperatures and timings carefully for best results.'
    )

    # --- Recipe 1: Classic Roast Chicken ---
    doc.add_heading('Classic Roast Chicken', level=1)
    doc.add_paragraph(
        'A perfectly roasted chicken is the cornerstone of home cooking. '
        'This recipe yields a golden, crispy skin with juicy meat inside.'
    )
    doc.add_heading('Ingredients', level=2)
    doc.add_paragraph('1 whole chicken (about 4 lbs)', style='List Bullet')
    doc.add_paragraph('2 tablespoons olive oil', style='List Bullet')
    doc.add_paragraph('1 lemon, halved', style='List Bullet')
    doc.add_paragraph('4 cloves garlic', style='List Bullet')
    doc.add_paragraph('Fresh rosemary and thyme sprigs', style='List Bullet')
    doc.add_paragraph('Salt and freshly ground black pepper', style='List Bullet')
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph(
        '1. Preheat your oven to 425\u00b0F. Pat the chicken dry with paper towels '
        'and rub with olive oil, salt, and pepper.'
    )
    doc.add_paragraph(
        '2. Stuff the cavity with lemon halves, garlic, and herbs. '
        'Tie the legs together with kitchen twine.'
    )
    doc.add_paragraph(
        '3. Place in a roasting pan and roast for 1 hour and 15 minutes, '
        'or until the internal temperature reaches 165\u00b0F at the thickest part of the thigh.'
    )
    doc.add_paragraph(
        '4. Let rest for 15 minutes before carving. Save the drippings for gravy.'
    )

    # --- Recipe 2: Homemade Sourdough Bread ---
    doc.add_heading('Homemade Sourdough Bread', level=1)
    doc.add_paragraph(
        'Nothing beats the aroma of fresh bread baking in the oven. '
        'This sourdough recipe requires patience but delivers incredible flavor.'
    )
    doc.add_heading('Ingredients', level=2)
    doc.add_paragraph('500g bread flour', style='List Bullet')
    doc.add_paragraph('350g water (room temperature)', style='List Bullet')
    doc.add_paragraph('100g active sourdough starter', style='List Bullet')
    doc.add_paragraph('10g sea salt', style='List Bullet')
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph(
        '1. Mix flour and water, let rest 30 minutes (autolyse). '
        'Add starter and salt, fold until combined.'
    )
    doc.add_paragraph(
        '2. Perform stretch and folds every 30 minutes for the first 2 hours. '
        'Then let bulk ferment for 4-6 hours at room temperature.'
    )
    doc.add_paragraph(
        '3. Shape the dough and place in a banneton. Refrigerate overnight for 12-16 hours.'
    )
    doc.add_paragraph(
        '4. Preheat your Dutch oven inside the oven to 450\u00b0F for at least 45 minutes. '
        'Score the dough and bake covered for 20 minutes, then uncovered for another 25 minutes.'
    )
    doc.add_paragraph(
        '5. The bread is done when the internal temperature reaches 200\u00b0F. '
        'Cool completely on a wire rack before slicing.'
    )

    # --- Recipe 3: Slow-Roasted Pulled Pork ---
    doc.add_heading('Slow-Roasted Pulled Pork', level=1)
    doc.add_paragraph(
        'This low-and-slow pulled pork is perfect for summer barbecue gatherings. '
        'The meat becomes incredibly tender and full of smoky flavor.'
    )
    doc.add_heading('Ingredients', level=2)
    doc.add_paragraph('1 bone-in pork shoulder (8-10 lbs)', style='List Bullet')
    doc.add_paragraph('3 tablespoons brown sugar', style='List Bullet')
    doc.add_paragraph('2 tablespoons smoked paprika', style='List Bullet')
    doc.add_paragraph('1 tablespoon garlic powder', style='List Bullet')
    doc.add_paragraph('1 tablespoon onion powder', style='List Bullet')
    doc.add_paragraph('1 cup apple cider vinegar', style='List Bullet')
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph(
        '1. Combine brown sugar, paprika, garlic powder, onion powder, salt, and pepper. '
        'Rub generously all over the pork shoulder.'
    )
    doc.add_paragraph(
        '2. Preheat the oven to 275\u00b0F. Place the pork in a large roasting pan '
        'and pour apple cider vinegar around it.'
    )
    doc.add_paragraph(
        '3. Cover tightly with aluminum foil and roast for 8-10 hours until the meat '
        'falls apart easily when pulled with a fork.'
    )
    doc.add_paragraph(
        '4. Shred the meat with two forks and toss with your favorite barbecue sauce.'
    )

    # --- Recipe 4: Crispy Pizza Margherita ---
    doc.add_heading('Crispy Pizza Margherita', level=1)
    doc.add_paragraph(
        'A classic Neapolitan-style pizza with a thin, crispy crust. '
        'The key is a very hot oven and simple, quality ingredients.'
    )
    doc.add_heading('Ingredients', level=2)
    doc.add_paragraph('300g "00" flour or bread flour', style='List Bullet')
    doc.add_paragraph('200ml warm water', style='List Bullet')
    doc.add_paragraph('7g instant yeast', style='List Bullet')
    doc.add_paragraph('1 teaspoon sugar', style='List Bullet')
    doc.add_paragraph('San Marzano tomato sauce', style='List Bullet')
    doc.add_paragraph('Fresh mozzarella, sliced', style='List Bullet')
    doc.add_paragraph('Fresh basil leaves', style='List Bullet')
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph(
        '1. Mix flour, yeast, sugar, salt, and water. Knead for 10 minutes until smooth. '
        'Let rise 1 hour.'
    )
    doc.add_paragraph(
        '2. Preheat your oven to 400\u00b0F with a pizza stone or inverted baking sheet inside.'
    )
    doc.add_paragraph(
        '3. Stretch the dough into a thin round. Add sauce, mozzarella, and a drizzle of olive oil.'
    )
    doc.add_paragraph(
        '4. Bake for 12-15 minutes until the crust is golden and the cheese is bubbly. '
        'Top with fresh basil after removing from the oven.'
    )

    # --- Recipe 5: Grandma\'s Apple Pie ---
    doc.add_heading("Grandma's Apple Pie", level=1)
    doc.add_paragraph(
        'This double-crust apple pie has been in our family for three generations. '
        'The secret is mixing tart Granny Smith apples with sweet Honeycrisp.'
    )
    doc.add_heading('Ingredients', level=2)
    doc.add_paragraph('6 large apples (mix of Granny Smith and Honeycrisp)', style='List Bullet')
    doc.add_paragraph('3/4 cup granulated sugar', style='List Bullet')
    doc.add_paragraph('2 tablespoons all-purpose flour', style='List Bullet')
    doc.add_paragraph('1 teaspoon cinnamon', style='List Bullet')
    doc.add_paragraph('1/4 teaspoon nutmeg', style='List Bullet')
    doc.add_paragraph('2 prepared pie crusts', style='List Bullet')
    doc.add_paragraph('2 tablespoons butter, cut into small pieces', style='List Bullet')
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph(
        '1. Preheat the oven to 375\u00b0F. Peel, core, and slice the apples thinly.'
    )
    doc.add_paragraph(
        '2. Toss the apples with sugar, flour, cinnamon, and nutmeg.'
    )
    doc.add_paragraph(
        '3. Line a 9-inch pie dish with one crust. Fill with apple mixture and dot with butter. '
        'Cover with second crust, crimp edges, and cut vents.'
    )
    doc.add_paragraph(
        '4. Bake for 45-50 minutes. If the edges brown too quickly, cover them with foil. '
        'Reduce temperature to 350\u00b0F for the last 15 minutes if the top is browning too fast.'
    )

    # --- Tips Section ---
    doc.add_heading('General Baking Tips', level=1)
    doc.add_paragraph(
        'Always preheat your oven fully before baking. An oven thermometer is a worthwhile '
        'investment to ensure accuracy. Many home ovens can be off by as much as 25 degrees.'
    )
    doc.add_paragraph(
        'When a recipe calls for room temperature ingredients, plan ahead. Butter and eggs '
        'should sit out for about 30-45 minutes before use.'
    )
    doc.add_paragraph(
        'For delicate pastries and meringues, a lower oven temperature helps '
        'prevent cracking and ensures even baking throughout.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
