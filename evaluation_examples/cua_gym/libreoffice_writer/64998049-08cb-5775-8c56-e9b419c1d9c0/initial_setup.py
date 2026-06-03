"""
Initial Setup: Document with front matter and chapters, all using Arabic page numbering
Task ID: writer_tech_071
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_071'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section):
    """Add a footer with a PAGE field code to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Page number field code: BEGIN
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    # Instruction
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    # END
    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def set_page_num_type(section, fmt='decimal', start=None):
    """Set the page number format on a section via XML.
    fmt: 'decimal' for Arabic (1,2,3), 'lowerRoman' for (i,ii,iii)
    """
    sectPr = section._sectPr
    # Remove existing pgNumType if any
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    attrs = {qn('w:fmt'): fmt}
    if start is not None:
        attrs[qn('w:start')] = str(start)
    pgNumType = sectPr.makeelement(qn('w:pgNumType'), attrs)
    sectPr.append(pgNumType)


def create_initial():
    doc = Document()

    # ===== SECTION 1: Front Matter (Arabic numbering - this is the initial "wrong" state) =====
    section1 = doc.sections[0]
    section1.page_width = Inches(8.5)
    section1.page_height = Inches(11)
    section1.left_margin = Inches(1.25)
    section1.right_margin = Inches(1.25)
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)

    # Set Arabic numbering starting at 1 (the "wrong" state for front matter)
    set_page_num_type(section1, fmt='decimal', start=1)

    # Add footer with page number
    add_page_number_footer(section1)

    # --- Table of Contents page ---
    h = doc.add_heading('Table of Contents', level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    toc_entries = [
        ('Preface', 'ii'),
        ('Chapter 1: Introduction to Cloud Architecture', '1'),
        ('Chapter 2: Microservices Design Patterns', '8'),
        ('Chapter 3: Container Orchestration with Kubernetes', '17'),
        ('Chapter 4: Observability and Monitoring', '28'),
        ('Chapter 5: Security Best Practices', '39'),
    ]
    for title, page in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}')
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        tab_run = p.add_run(f'\t{page}')
        tab_run.font.size = Pt(12)
        tab_run.font.name = 'Calibri'

    doc.add_paragraph()  # spacing

    # --- Preface page (still in section 1 = front matter) ---
    # Add page break for preface
    doc.add_page_break()

    h2 = doc.add_heading('Preface', level=1)
    h2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    preface_text = [
        "This technical guide provides a comprehensive overview of modern cloud-native "
        "application architecture. Drawing from years of practical experience deploying "
        "production systems at scale, this book aims to bridge the gap between theoretical "
        "distributed systems concepts and real-world engineering practice.",

        "The rapid evolution of container technologies, orchestration platforms, and "
        "serverless computing has fundamentally changed how we design, build, and operate "
        "software systems. Organizations that once managed monolithic applications on "
        "dedicated hardware now face the challenge of coordinating hundreds of loosely "
        "coupled services across dynamic infrastructure.",

        "Each chapter in this guide addresses a specific aspect of cloud-native architecture, "
        "from the foundational principles of microservices decomposition to advanced topics "
        "like distributed tracing and chaos engineering. Code examples are provided in Python "
        "and Go, with Terraform configurations for infrastructure provisioning.",

        "We would like to thank the engineering teams at Meridian Systems, Apex Cloud Solutions, "
        "and the open-source contributors who reviewed early drafts and provided invaluable "
        "feedback on the technical accuracy of the material presented here.",
    ]
    for text in preface_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # ===== SECTION 2: Body / Chapters (new section, also Arabic - same "wrong" continuous numbering) =====
    # Add section break before Chapter 1
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section2 = doc.sections[1]
    section2.page_width = Inches(8.5)
    section2.page_height = Inches(11)
    section2.left_margin = Inches(1.25)
    section2.right_margin = Inches(1.25)
    section2.top_margin = Inches(1)
    section2.bottom_margin = Inches(1)

    # Arabic numbering, continuing (no restart) - "wrong" state
    set_page_num_type(section2, fmt='decimal')

    # Add footer with page number to section 2
    add_page_number_footer(section2)

    # --- Chapter 1: Introduction to Cloud Architecture ---
    h3 = doc.add_heading('Chapter 1: Introduction to Cloud Architecture', level=1)

    ch1_paragraphs = [
        "Cloud architecture refers to the components and subcomponents required for cloud "
        "computing. These components typically consist of a front-end platform, back-end "
        "platforms, a cloud-based delivery model, and a network. Combined, these components "
        "make up cloud computing architecture.",

        "The shift from monolithic to distributed architectures has been driven by the need "
        "for greater scalability, resilience, and deployment velocity. Modern applications "
        "must handle unpredictable traffic patterns, maintain sub-second response times, "
        "and support continuous deployment cycles measured in hours rather than months.",

        "1.1 The Twelve-Factor App Methodology",

        "The twelve-factor app methodology, originally published by engineers at Heroku in "
        "2012, established foundational principles for building software-as-a-service "
        "applications. These factors address codebase management, dependency isolation, "
        "configuration externalization, backing service abstraction, build-release-run "
        "separation, and process statefulness.",

        "Factor I (Codebase) mandates that each deployable service maintains exactly one "
        "codebase tracked in version control, with many deploys across staging and production "
        "environments. Factor III (Config) requires strict separation of configuration from "
        "code, storing environment-specific values in environment variables rather than "
        "hardcoded constants or checked-in configuration files.",

        "1.2 Service Mesh Architecture",

        "A service mesh provides a dedicated infrastructure layer for handling service-to-service "
        "communication. Implementations such as Istio, Linkerd, and Consul Connect deploy "
        "sidecar proxies alongside each service instance, enabling features like mutual TLS "
        "authentication, load balancing, circuit breaking, and distributed tracing without "
        "requiring application-level code changes.",
    ]
    for text in ch1_paragraphs:
        if text.startswith('1.'):
            h = doc.add_heading(text, level=2)
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

    # --- Chapter 2: Microservices Design Patterns ---
    doc.add_page_break()
    h4 = doc.add_heading('Chapter 2: Microservices Design Patterns', level=1)

    ch2_paragraphs = [
        "Microservices architecture decomposes applications into small, independently "
        "deployable services organized around business capabilities. Each service owns "
        "its data store and communicates with other services through well-defined APIs, "
        "typically REST or gRPC.",

        "2.1 Saga Pattern for Distributed Transactions",

        "The saga pattern manages data consistency across microservices by coordinating "
        "a sequence of local transactions. Each service performs its local transaction and "
        "publishes an event that triggers the next step. If any step fails, compensating "
        "transactions are executed to undo prior changes.",

        "Consider an e-commerce order flow: the Order Service creates an order, the "
        "Payment Service charges the customer, and the Inventory Service reserves stock. "
        "If inventory reservation fails, the saga orchestrator triggers a payment refund "
        "and order cancellation, maintaining eventual consistency without distributed locks.",

        "2.2 CQRS and Event Sourcing",

        "Command Query Responsibility Segregation (CQRS) separates read and write models, "
        "allowing independent optimization of each path. Write operations process commands "
        "through an aggregate root that enforces business invariants, while read operations "
        "query denormalized projections optimized for specific UI requirements.",

        "Event sourcing complements CQRS by persisting the full history of state changes "
        "as an append-only log of domain events. The current state is reconstructed by "
        "replaying events, providing a complete audit trail and enabling temporal queries "
        "that answer questions about past system states.",
    ]
    for text in ch2_paragraphs:
        if text.startswith('2.'):
            h = doc.add_heading(text, level=2)
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

    # --- Chapter 3: Container Orchestration ---
    doc.add_page_break()
    h5 = doc.add_heading('Chapter 3: Container Orchestration with Kubernetes', level=1)

    ch3_paragraphs = [
        "Kubernetes has emerged as the de facto standard for container orchestration, "
        "providing declarative configuration, automated rollouts, self-healing capabilities, "
        "and horizontal scaling for containerized workloads. Understanding its architecture "
        "and operational patterns is essential for modern platform engineering teams.",

        "3.1 Pod Scheduling and Resource Management",

        "The Kubernetes scheduler assigns pods to nodes based on resource requests, limits, "
        "affinity rules, and topology constraints. Resource requests guarantee minimum CPU "
        "and memory allocation, while limits cap maximum consumption to prevent noisy-neighbor "
        "effects in multi-tenant clusters.",

        "Production deployments should define resource requests and limits for every container. "
        "A common configuration for a Java microservice might specify 500m CPU request with "
        "1000m limit, and 512Mi memory request with 1Gi limit, adjusted based on profiling "
        "data from load testing environments.",
    ]
    for text in ch3_paragraphs:
        if text.startswith('3.'):
            h = doc.add_heading(text, level=2)
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
