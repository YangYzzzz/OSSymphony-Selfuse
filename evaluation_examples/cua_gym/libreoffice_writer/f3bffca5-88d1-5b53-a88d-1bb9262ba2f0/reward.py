"""
Reward Script: Probationary Period Review Form
Task ID: writer_hr_084
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25) - Employee Information section with fields
  Component 2 (0.35) - Competency evaluation table with 6 areas, ratings, and comments
  Component 3 (0.20) - Recommendation section with 3 options
  Component 4 (0.20) - Signature blocks for Manager and HR
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_084'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts for searching
    all_para_texts = [p.text.strip() for p in doc.paragraphs]
    all_para_lower = [t.lower() for t in all_para_texts]

    # =========================================================================
    # Component 1: Employee Information Section (0.25 points)
    # The golden doc has a "Employee Information" heading and fields like
    # Employee Name, Department, Position/Title, Start Date, Review Date,
    # Reporting Manager.
    # Initial doc has NONE of this - only the title.
    # =========================================================================
    try:
        # Check for Employee Information heading
        has_emp_info_heading = False
        for p in doc.paragraphs:
            if 'employee information' in p.text.strip().lower():
                has_emp_info_heading = True
                break

        # Check for employee info fields (at least 4 of 6 expected)
        expected_fields = ['employee name', 'department', 'position', 'start date', 'review date', 'manager']
        found_fields = 0
        for field in expected_fields:
            for t in all_para_lower:
                if field in t:
                    found_fields += 1
                    break

        if has_emp_info_heading and found_fields >= 4:
            print(f"PASS: Component 1 - Employee Information section found with heading and {found_fields}/6 fields (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - Employee Information: heading={has_emp_info_heading}, fields={found_fields}/6 (need heading + >=4 fields)")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # =========================================================================
    # Component 2: Competency Evaluation Table (0.35 points)
    # Golden has a table with 6 competency areas (Job Knowledge, Quality of Work,
    # Attendance, Communication, Initiative, Teamwork), rating columns
    # (Excellent, Satisfactory, Needs Improvement, Unsatisfactory), and
    # comment rows for each competency.
    # Initial doc has NO tables at all.
    # =========================================================================
    try:
        competency_areas = ['job knowledge', 'quality of work', 'attendance', 'communication', 'initiative', 'teamwork']
        rating_columns = ['excellent', 'satisfactory', 'needs improvement', 'unsatisfactory']

        # Find the competency evaluation table
        comp_table_found = False
        competencies_found = 0
        ratings_found = 0
        comment_rows_found = 0

        for table in doc.tables:
            # Collect all cell texts in this table
            table_cell_texts = []
            for row in table.rows:
                row_texts = [cell.text.strip().lower() for cell in row.cells]
                table_cell_texts.append(row_texts)

            # Check if header row has rating columns
            if len(table.rows) > 0:
                header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
                header_joined = ' '.join(header_texts)
                ratings_in_header = sum(1 for r in rating_columns if r in header_joined)

                if ratings_in_header >= 3:
                    comp_table_found = True
                    ratings_found = ratings_in_header

                    # Check for competency areas in first column
                    for row in table.rows:
                        first_cell = row.cells[0].text.strip().lower()
                        for comp in competency_areas:
                            if comp in first_cell and 'comment' not in first_cell:
                                competencies_found += 1
                                break

                    # Check for comment rows
                    for row in table.rows:
                        first_cell = row.cells[0].text.strip().lower()
                        if 'comment' in first_cell:
                            comment_rows_found += 1

                    break  # Found the competency table

        sub_score = 0.0
        if comp_table_found:
            sub_score += 0.10  # Table with rating headers exists
        if competencies_found >= 5:
            sub_score += 0.15  # At least 5 of 6 competency areas present
        elif competencies_found >= 3:
            sub_score += 0.08  # Partial credit
        if comment_rows_found >= 4:
            sub_score += 0.10  # Comment rows for competencies
        elif comment_rows_found >= 2:
            sub_score += 0.05  # Partial credit

        if sub_score > 0:
            print(f"PASS: Component 2 - Competency table: table={comp_table_found}, competencies={competencies_found}/6, ratings={ratings_found}/4, comments={comment_rows_found} ({sub_score} pts)")
            total_score += sub_score
        else:
            print(f"FAIL: Component 2 - No competency evaluation table found with required structure")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # =========================================================================
    # Component 3: Recommendation Section (0.20 points)
    # Golden has a Recommendation heading and three options:
    # - Confirm Employment
    # - Extend Probation (with specify period)
    # - Terminate Employment
    # Initial doc has NONE of this.
    # =========================================================================
    try:
        has_recommendation_heading = False
        for p in doc.paragraphs:
            if 'recommendation' in p.text.strip().lower():
                has_recommendation_heading = True
                break

        options_found = 0
        option_keywords = ['confirm employment', 'extend probation', 'terminate']
        for kw in option_keywords:
            for t in all_para_lower:
                if kw in t:
                    options_found += 1
                    break

        if has_recommendation_heading and options_found >= 2:
            print(f"PASS: Component 3 - Recommendation section with heading and {options_found}/3 options (0.20 pts)")
            total_score += 0.20
        elif has_recommendation_heading and options_found >= 1:
            print(f"PARTIAL: Component 3 - Recommendation heading found but only {options_found}/3 options (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 - Recommendation: heading={has_recommendation_heading}, options={options_found}/3")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # =========================================================================
    # Component 4: Signature Blocks for Manager and HR (0.20 points)
    # Golden has a "Signatures" heading and a 2-column table with Manager and
    # Human Resources columns, each with Signature, Name, Date fields.
    # Initial doc has NONE of this.
    # =========================================================================
    try:
        has_signature_heading = False
        for p in doc.paragraphs:
            if 'signature' in p.text.strip().lower():
                has_signature_heading = True
                break

        # Look for a signature table with Manager and HR columns
        sig_table_found = False
        has_manager_col = False
        has_hr_col = False
        has_sig_fields = False

        for table in doc.tables:
            table_text_all = ' '.join(cell.text.strip().lower() for row in table.rows for cell in row.cells)
            if 'manager' in table_text_all and ('human resources' in table_text_all or 'hr' in table_text_all):
                sig_table_found = True
                has_manager_col = True
                has_hr_col = True
                # Check for signature/name/date fields
                if 'signature' in table_text_all:
                    has_sig_fields = True
                break

        # Also check paragraphs for non-table signature blocks
        if not sig_table_found:
            manager_sig = False
            hr_sig = False
            for t in all_para_lower:
                if 'manager' in t and 'signature' in t:
                    manager_sig = True
                if ('human resources' in t or 'hr' in t) and 'signature' in t:
                    hr_sig = True
            if manager_sig and hr_sig:
                sig_table_found = True
                has_manager_col = True
                has_hr_col = True
                has_sig_fields = True

        sub_score_4 = 0.0
        if has_signature_heading or sig_table_found:
            if sig_table_found and has_sig_fields:
                sub_score_4 = 0.20
                print(f"PASS: Component 4 - Signature blocks with Manager and HR ({sub_score_4} pts)")
            elif sig_table_found:
                sub_score_4 = 0.15
                print(f"PARTIAL: Component 4 - Signature table found but missing some fields ({sub_score_4} pts)")
            else:
                sub_score_4 = 0.05
                print(f"PARTIAL: Component 4 - Signature heading found but no structured blocks ({sub_score_4} pts)")
            total_score += sub_score_4
        else:
            print(f"FAIL: Component 4 - No signature blocks found")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
