"""
Initial Setup: Create a technical report with headings at levels 1-4, no TOC.
Task ID: writer_mt_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_052'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Page 1: Reserved for TOC (blank page with just a title) ---
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = toc_title.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(18)

    # Add empty lines to indicate TOC placeholder area
    for _ in range(3):
        doc.add_paragraph("")

    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    note_run = toc_note.add_run("[TOC to be inserted here]")
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Page break after TOC placeholder page
    doc.add_page_break()

    # =================================================================
    # SECTION 1: Introduction (Heading 1)
    # =================================================================
    doc.add_heading("1. Introduction", level=1)

    # H2 entries under Introduction
    doc.add_heading("1.1 Background and Motivation", level=2)
    doc.add_paragraph(
        "The rapid advancement of distributed computing systems has created new "
        "challenges in ensuring data consistency across geographically dispersed "
        "nodes. This report examines the current state of fault-tolerant consensus "
        "protocols and proposes a novel approach to reducing latency in wide-area "
        "network deployments."
    )

    doc.add_heading("1.1.1 Historical Context", level=3)
    doc.add_paragraph(
        "Early distributed systems relied on simple two-phase commit protocols, "
        "which proved inadequate for modern cloud-scale applications. The seminal "
        "work by Lamport on Paxos in 1989 established the theoretical foundations "
        "for consensus in asynchronous systems."
    )

    doc.add_heading("1.1.1.1 Pre-Internet Era Systems", level=4)
    doc.add_paragraph(
        "Mainframe-based transaction processing systems of the 1970s and 1980s "
        "used centralized coordination models that could not scale beyond a single "
        "data center. Notable examples include IBM's CICS and Tandem's NonStop."
    )

    doc.add_heading("1.1.1.2 Early Network Protocols", level=4)
    doc.add_paragraph(
        "The emergence of TCP/IP and local area networks in the late 1980s enabled "
        "the first generation of truly distributed databases, including Oracle RAC "
        "and Microsoft SQL Server clustering solutions."
    )

    doc.add_heading("1.1.2 Current Industry Landscape", level=3)
    doc.add_paragraph(
        "Modern cloud providers such as Amazon Web Services, Google Cloud Platform, "
        "and Microsoft Azure each implement proprietary consensus mechanisms. Google's "
        "Spanner uses TrueTime for global consistency, while Amazon's DynamoDB employs "
        "vector clocks for conflict resolution."
    )

    doc.add_heading("1.2 Research Objectives", level=2)
    doc.add_paragraph(
        "This study aims to: (1) characterize the performance bottlenecks in existing "
        "consensus protocols under realistic network conditions, (2) develop an adaptive "
        "protocol selection framework, and (3) validate the approach through extensive "
        "simulation and testbed experiments."
    )

    doc.add_heading("1.2.1 Scope and Limitations", level=3)
    doc.add_paragraph(
        "The research focuses on Byzantine fault tolerance in permissioned blockchain "
        "networks with up to 100 participating nodes. We do not address permissionless "
        "systems or networks exceeding 1000 nodes."
    )

    doc.add_heading("1.2.1.1 Threat Model Assumptions", level=4)
    doc.add_paragraph(
        "We assume a partially synchronous network model where message delays are "
        "bounded but unknown. Up to f = (n-1)/3 nodes may exhibit Byzantine behavior, "
        "including equivocation and selective message dropping."
    )

    doc.add_heading("1.2.2 Key Contributions", level=3)
    doc.add_paragraph(
        "The primary contributions of this work include a formal proof of liveness "
        "under partial synchrony, an optimized message complexity of O(n log n) per "
        "consensus round, and a reference implementation in Rust with comprehensive "
        "benchmarks."
    )

    # =================================================================
    # SECTION 2: Methodology (Heading 1)
    # =================================================================
    doc.add_heading("2. Methodology", level=1)

    doc.add_heading("2.1 Experimental Design", level=2)
    doc.add_paragraph(
        "Our experimental methodology follows a three-phase approach: controlled "
        "simulation, emulated network testing, and live deployment validation. Each "
        "phase builds upon the results of the previous one, progressively increasing "
        "the realism of the evaluation environment."
    )

    doc.add_heading("2.1.1 Simulation Framework", level=3)
    doc.add_paragraph(
        "We developed a discrete-event simulator capable of modeling network partitions, "
        "variable latency distributions, and node failures. The simulator supports "
        "pluggable consensus protocol implementations through a standardized interface."
    )

    doc.add_heading("2.1.1.1 Network Topology Models", level=4)
    doc.add_paragraph(
        "Three topology models were evaluated: fully connected mesh, hierarchical "
        "cluster-based, and geographic region-aware. Each model was parameterized "
        "with latency distributions derived from real-world measurements collected "
        "from AWS CloudPing data over a 6-month period."
    )

    doc.add_heading("2.1.1.2 Failure Injection Strategies", level=4)
    doc.add_paragraph(
        "Byzantine failures were injected using a Markov chain model with configurable "
        "transition probabilities. Crash failures followed an exponential distribution "
        "with a mean time between failures (MTBF) of 24 hours per node."
    )

    doc.add_heading("2.1.2 Emulation Testbed", level=3)
    doc.add_paragraph(
        "The emulation testbed consisted of 50 Docker containers deployed across "
        "5 physical servers connected via a software-defined network. Traffic shaping "
        "was applied using Linux tc to replicate inter-continental latency profiles."
    )

    doc.add_heading("2.1.3 Production Pilot Deployment", level=3)
    doc.add_paragraph(
        "A 30-day pilot deployment was conducted on a consortium blockchain network "
        "spanning data centers in Frankfurt, Singapore, and Virginia. The network "
        "processed approximately 2,500 transactions per second during peak hours."
    )

    doc.add_heading("2.2 Data Collection and Analysis", level=2)
    doc.add_paragraph(
        "Performance metrics were collected using a custom telemetry pipeline built "
        "on Apache Kafka and InfluxDB. Key metrics include consensus latency (p50, "
        "p95, p99), throughput (transactions per second), and message overhead "
        "(bytes per consensus round)."
    )

    doc.add_heading("2.2.1 Statistical Methods", level=3)
    doc.add_paragraph(
        "All performance comparisons use the Mann-Whitney U test with Bonferroni "
        "correction for multiple comparisons. Effect sizes are reported using "
        "Cliff's delta. Confidence intervals are computed at the 95% level using "
        "bootstrap resampling with 10,000 iterations."
    )

    doc.add_heading("2.2.1.1 Outlier Detection", level=4)
    doc.add_paragraph(
        "Outliers in latency measurements were identified using the modified Z-score "
        "method with a threshold of 3.5. Approximately 2.3% of measurements were "
        "classified as outliers and excluded from aggregate statistics."
    )

    doc.add_heading("2.2.2 Reproducibility Framework", level=3)
    doc.add_paragraph(
        "All experiments are fully reproducible using the provided Docker Compose "
        "configuration and Ansible playbooks. Random seeds are fixed and documented "
        "for each experimental run. The complete dataset and analysis scripts are "
        "available in the supplementary materials."
    )

    # =================================================================
    # SECTION 3: Results and Discussion (Heading 1)
    # =================================================================
    doc.add_heading("3. Results and Discussion", level=1)

    doc.add_heading("3.1 Performance Benchmarks", level=2)
    doc.add_paragraph(
        "The adaptive protocol selection framework achieved a 37% reduction in "
        "median consensus latency compared to static Raft deployment, and a 52% "
        "reduction compared to PBFT in high-contention scenarios. Throughput "
        "improvements ranged from 18% to 43% depending on network conditions."
    )

    doc.add_heading("3.1.1 Latency Analysis", level=3)
    doc.add_paragraph(
        "Under normal operating conditions (no failures, symmetric latency), the "
        "adaptive framework selected the optimistic fast path in 94.7% of consensus "
        "rounds. The median round-trip latency was 12.3ms for intra-region consensus "
        "and 187.5ms for cross-region consensus."
    )

    doc.add_heading("3.1.1.1 Tail Latency Characteristics", level=4)
    doc.add_paragraph(
        "The p99 latency showed the most significant improvement, decreasing from "
        "2,340ms (static PBFT) to 456ms (adaptive framework). This 5.1x improvement "
        "is attributed to the framework's ability to bypass the view-change protocol "
        "in cases of transient network delays."
    )

    doc.add_heading("3.1.2 Throughput Scaling", level=3)
    doc.add_paragraph(
        "Throughput scaled linearly up to 50 nodes, after which coordination overhead "
        "caused sub-linear growth. At 100 nodes, the system sustained 8,200 TPS for "
        "read-heavy workloads and 3,100 TPS for write-intensive scenarios."
    )

    doc.add_heading("3.2 Fault Tolerance Evaluation", level=2)
    doc.add_paragraph(
        "The system maintained consistency guarantees under all tested failure "
        "scenarios, including simultaneous crash failures of up to 32 out of 100 "
        "nodes. Recovery time after network partition healing averaged 4.7 seconds "
        "with a standard deviation of 1.2 seconds."
    )

    doc.add_heading("3.2.1 Byzantine Fault Scenarios", level=3)
    doc.add_paragraph(
        "In scenarios with active Byzantine adversaries controlling up to 33% of "
        "nodes, the system correctly identified and isolated malicious participants "
        "within an average of 3 consensus rounds. No safety violations were observed "
        "across 10 million simulated transactions."
    )

    doc.add_heading("3.2.2 Recovery Performance", level=3)
    doc.add_paragraph(
        "State synchronization after node recovery completed in under 30 seconds "
        "for ledgers containing up to 1 million transactions. The incremental "
        "snapshot transfer protocol reduced bandwidth consumption by 78% compared "
        "to full state transfer."
    )

    doc.add_heading("3.2.2.1 Checkpoint Optimization", level=4)
    doc.add_paragraph(
        "The adaptive checkpoint interval algorithm reduced storage overhead by "
        "45% while maintaining a maximum recovery time of 60 seconds. Checkpoints "
        "are triggered based on a combination of elapsed time (30-second minimum) "
        "and accumulated state changes (threshold: 10,000 operations)."
    )

    doc.add_heading("3.3 Comparative Analysis", level=2)
    doc.add_paragraph(
        "A comprehensive comparison with six existing consensus protocols reveals "
        "that the adaptive framework consistently outperforms static deployments "
        "across all evaluated metrics. The most significant gains are observed in "
        "heterogeneous network environments with variable latency profiles."
    )

    doc.add_heading("3.3.1 Protocol Selection Accuracy", level=3)
    doc.add_paragraph(
        "The machine learning-based protocol selector achieved 91.2% accuracy in "
        "choosing the optimal protocol for a given network condition. Misclassifications "
        "predominantly occurred during rapid network state transitions, with a median "
        "recovery time of 2.1 seconds to correct the selection."
    )

    doc.add_heading("3.3.1.1 Feature Importance Analysis", level=4)
    doc.add_paragraph(
        "The top three features for protocol selection were network latency variance "
        "(importance: 0.34), active node count (importance: 0.22), and recent failure "
        "rate (importance: 0.18). Geographic distribution of nodes contributed 0.11 "
        "to the overall prediction accuracy."
    )

    doc.add_heading("3.4 Security Implications", level=2)
    doc.add_paragraph(
        "The adaptive protocol switching mechanism introduces potential attack vectors "
        "that must be carefully mitigated. An adversary could attempt to manipulate "
        "network conditions to force selection of a less secure consensus protocol."
    )

    doc.add_heading("3.4.1 Attack Surface Analysis", level=3)
    doc.add_paragraph(
        "We identified three primary attack vectors: forced protocol downgrade through "
        "latency injection, selector poisoning via fabricated telemetry data, and "
        "timing-based side-channel attacks during protocol transitions. Mitigation "
        "strategies include minimum security thresholds and anomaly detection."
    )

    doc.add_heading("3.4.1.1 Mitigation Effectiveness", level=4)
    doc.add_paragraph(
        "The implemented countermeasures successfully prevented all tested attack "
        "scenarios with less than 3% overhead on normal operation throughput. The "
        "anomaly detection system exhibited a false positive rate of 0.7% and a "
        "false negative rate of 0.2% across 50,000 simulated attack attempts."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Verify heading counts
    verify_doc = Document(OUTPUT)
    h1 = h2 = h3 = h4 = 0
    for p in verify_doc.paragraphs:
        if p.style.name == 'Heading 1':
            h1 += 1
        elif p.style.name == 'Heading 2':
            h2 += 1
        elif p.style.name == 'Heading 3':
            h3 += 1
        elif p.style.name == 'Heading 4':
            h4 += 1
    print(f'Heading counts: H1={h1}, H2={h2}, H3={h3}, H4={h4}, Total={h1+h2+h3+h4}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
