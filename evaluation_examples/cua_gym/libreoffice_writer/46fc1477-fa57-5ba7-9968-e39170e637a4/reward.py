"""
Reward Script: Create a table of contents for a report with headings levels 1-4
Task ID: writer_mt_052
Domain: libreoffice_writer
Scoring:
  Precondition gate: Original headings preserved (not scored — true in both envs)
  Component 1 (0.5): TOC entries exist — 36 entries with heading text and page numbers
  Component 2 (0.3): Correct indentation hierarchy across 4 levels
  Component 3 (0.2): Dot leader tab stops on TOC entries
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_052'

# Expected heading texts extracted from the task context (36 total)
EXPECTED_HEADINGS = {
    1: [  # Heading 1 — 3 entries
        '1. Introduction',
        '2. Methodology',
        '3. Results and Discussion',
    ],
    2: [  # Heading 2 — 8 entries
        '1.1 Background and Motivation',
        '1.2 Research Objectives',
        '2.1 Experimental Design',
        '2.2 Data Collection and Analysis',
        '3.1 Performance Benchmarks',
        '3.2 Fault Tolerance Evaluation',
        '3.3 Comparative Analysis',
        '3.4 Security Implications',
    ],
    3: [  # Heading 3 — 15 entries
        '1.1.1 Historical Context',
        '1.1.2 Current Industry Landscape',
        '1.2.1 Scope and Limitations',
        '1.2.2 Key Contributions',
        '2.1.1 Simulation Framework',
        '2.1.2 Emulation Testbed',
        '2.1.3 Production Pilot Deployment',
        '2.2.1 Statistical Methods',
        '2.2.2 Reproducibility Framework',
        '3.1.1 Latency Analysis',
        '3.1.2 Throughput Scaling',
        '3.2.1 Byzantine Fault Scenarios',
        '3.2.2 Recovery Performance',
        '3.3.1 Protocol Selection Accuracy',
        '3.4.1 Attack Surface Analysis',
    ],
    4: [  # Heading 4 — 10 entries
        '1.1.1.1 Pre-Internet Era Systems',
        '1.1.1.2 Early Network Protocols',
        '1.2.1.1 Threat Model Assumptions',
        '2.1.1.1 Network Topology Models',
        '2.1.1.2 Failure Injection Strategies',
        '2.2.1.1 Outlier Detection',
        '3.1.1.1 Tail Latency Characteristics',
        '3.2.2.1 Checkpoint Optimization',
        '3.3.1.1 Feature Importance Analysis',
        '3.4.1.1 Mitigation Effectiveness',
    ],
}

ALL_HEADING_TEXTS = []
for level in [1, 2, 3, 4]:
    ALL_HEADING_TEXTS.extend(EXPECTED_HEADINGS[level])


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Identify TOC entries: paragraphs that contain heading text followed by
    # a tab and a page number, positioned BEFORE the actual heading paragraphs.
    # A TOC entry pattern: "<heading text>\t<page_number>"
    toc_entries = []  # list of (para_index, heading_text, page_num_str)
    first_heading_index = None

    for i, para in enumerate(doc.paragraphs):
        if para.style and 'Heading' in para.style.name:
            first_heading_index = i
            break

    if first_heading_index is None:
        first_heading_index = len(doc.paragraphs)

    # Scan paragraphs before the first heading for TOC-like entries
    for i in range(first_heading_index):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue
        # Check if this looks like a TOC entry: text + tab + number
        match = re.match(r'^(.+?)\t(\d+)\s*$', text)
        if match:
            heading_text = match.group(1).strip()
            page_num = match.group(2)
            toc_entries.append((i, heading_text, page_num))

    print(f"INFO: Found {len(toc_entries)} TOC entries before first heading at para {first_heading_index}")

    # Precondition gate: headings preserved (not scored — true in both initial and golden)
    try:
        heading_styles = {}
        for para in doc.paragraphs:
            if para.style and para.style.name in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'):
                level_name = para.style.name
                heading_styles[level_name] = heading_styles.get(level_name, 0) + 1

        h1 = heading_styles.get('Heading 1', 0)
        h2 = heading_styles.get('Heading 2', 0)
        h3 = heading_styles.get('Heading 3', 0)
        h4 = heading_styles.get('Heading 4', 0)

        if h1 < 3 or h2 < 8 or h3 < 15 or h4 < 10:
            print(f"GATE FAIL: Headings corrupted: H1={h1}, H2={h2}, H3={h3}, H4={h4} — returning 0.0")
            print("REWARD: 0.0")
            return 0.0
        else:
            print(f"GATE PASS: Headings preserved: H1={h1}, H2={h2}, H3={h3}, H4={h4}")
    except Exception as e:
        print(f"ERROR: Heading gate check — {e}")

    # Component 1: TOC entries exist — at least 30 of 36 expected entries (0.5 points)
    # We check that each expected heading appears as a TOC entry
    try:
        found_entries = set()
        for _, heading_text, _ in toc_entries:
            for expected in ALL_HEADING_TEXTS:
                if expected.lower() == heading_text.lower():
                    found_entries.add(expected)
                    break

        match_count = len(found_entries)
        total_expected = len(ALL_HEADING_TEXTS)  # 36

        if match_count >= 34:
            comp1_score = 0.5
            print(f"PASS: Component 1 — {match_count}/{total_expected} heading entries found in TOC (0.5 pts)")
        elif match_count >= 25:
            comp1_score = 0.25
            print(f"PARTIAL: Component 1 — {match_count}/{total_expected} heading entries found in TOC (0.25 pts)")
        else:
            comp1_score = 0.0
            print(f"FAIL: Component 1 — only {match_count}/{total_expected} heading entries found in TOC")
            if match_count < 5:
                missing = [h for h in ALL_HEADING_TEXTS if h not in found_entries]
                print(f"  Missing examples: {missing[:5]}")

        total_score += comp1_score
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Correct indentation hierarchy (0.3 points)
    # Level 1 = 0 indent, Level 2 = ~0.25in, Level 3 = ~0.50in, Level 4 = ~0.75in
    # We check the indent of TOC entries that match known headings at each level
    try:
        # Map found TOC entries to their heading level
        entry_indents = {}  # heading_text -> left_indent in EMU
        for idx, heading_text, _ in toc_entries:
            para = doc.paragraphs[idx]
            left_indent = para.paragraph_format.left_indent
            indent_emu = left_indent if left_indent is not None else 0
            entry_indents[heading_text.lower()] = indent_emu

        # Check indentation per level
        correct_levels = 0
        total_checked = 0

        # Expected indent values (EMU): 0, 228600, 457200, 685800
        indent_targets = {1: 0, 2: 228600, 3: 457200, 4: 685800}
        # Tolerance: 50000 EMU (~0.055 inches)
        tolerance = 50000

        for level in [1, 2, 3, 4]:
            for heading in EXPECTED_HEADINGS[level]:
                key = heading.lower()
                if key in entry_indents:
                    total_checked += 1
                    actual = entry_indents[key]
                    expected = indent_targets[level]
                    if abs(actual - expected) <= tolerance:
                        correct_levels += 1

        if total_checked > 0:
            indent_ratio = correct_levels / total_checked
        else:
            indent_ratio = 0.0

        if indent_ratio >= 0.9:
            comp2_score = 0.3
            print(f"PASS: Component 2 — {correct_levels}/{total_checked} entries have correct indentation (0.3 pts)")
        elif indent_ratio >= 0.5:
            comp2_score = 0.15
            print(f"PARTIAL: Component 2 — {correct_levels}/{total_checked} entries have correct indentation (0.15 pts)")
        else:
            comp2_score = 0.0
            print(f"FAIL: Component 2 — {correct_levels}/{total_checked} entries have correct indentation")

        total_score += comp2_score
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Dot leader tab stops on TOC entries (0.2 points)
    # Each TOC entry should have a RIGHT-aligned tab stop with DOTS leader
    try:
        entries_with_dots = 0
        entries_checked = 0

        for idx, heading_text, _ in toc_entries:
            para = doc.paragraphs[idx]
            entries_checked += 1
            for ts in para.paragraph_format.tab_stops:
                # Check for RIGHT alignment (value 2) and DOTS leader (value 1)
                if str(ts.alignment) == 'RIGHT (2)' and str(ts.leader) == 'DOTS (1)':
                    entries_with_dots += 1
                    break

        if entries_checked > 0:
            dot_ratio = entries_with_dots / entries_checked
        else:
            dot_ratio = 0.0

        if dot_ratio >= 0.9:
            comp3_score = 0.2
            print(f"PASS: Component 3 — {entries_with_dots}/{entries_checked} TOC entries have dot leaders (0.2 pts)")
        elif dot_ratio >= 0.5:
            comp3_score = 0.1
            print(f"PARTIAL: Component 3 — {entries_with_dots}/{entries_checked} TOC entries have dot leaders (0.1 pts)")
        else:
            comp3_score = 0.0
            print(f"FAIL: Component 3 — {entries_with_dots}/{entries_checked} TOC entries have dot leaders")

        total_score += comp3_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
