"""
Initial Setup: Church bulletin document with single-column layout
Task ID: writer_page_017
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'church_bulletin'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page Setup: Letter size, portrait, margins 1.5cm all sides, single column ---
    section = doc.sections[0]
    section.page_width = Cm(21.59)   # Letter width: 8.5 inches
    section.page_height = Cm(27.94)  # Letter height: 11 inches
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Ensure single column layout (default, but set explicitly)
    # Remove any existing cols element and add single-column one
    sectPr = section._sectPr
    # Remove existing cols if any
    for cols_elem in sectPr.findall(qn('w:cols')):
        sectPr.remove(cols_elem)
    # Add explicit single column
    cols = etree.SubElement(sectPr, qn('w:cols'))
    cols.set(qn('w:num'), '1')

    # --- Page 1: Church Bulletin ---
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('ST. MARK\'S COMMUNITY CHURCH')
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('Sunday Bulletin — March 9, 2025')
    subtitle_run.font.size = Pt(12)

    doc.add_paragraph()  # Spacer

    # Welcome section
    welcome_heading = doc.add_paragraph()
    welcome_run = welcome_heading.add_run('WELCOME & ANNOUNCEMENTS')
    welcome_run.bold = True
    welcome_run.font.size = Pt(11)

    doc.add_paragraph(
        'Welcome to St. Mark\'s Community Church! We are delighted to have you '
        'worship with us today. Whether you are a longtime member or a first-time '
        'visitor, we hope you feel at home in our congregation.'
    )

    doc.add_paragraph(
        'Our mission is to serve God and our community through worship, fellowship, '
        'education, and outreach. We invite you to join us in all our ministries.'
    )

    # Service Order section
    service_heading = doc.add_paragraph()
    service_run = service_heading.add_run('ORDER OF WORSHIP')
    service_run.bold = True
    service_run.font.size = Pt(11)

    service_items = [
        ('Prelude', 'Organ Music — "Morning Has Broken"'),
        ('Call to Worship', 'Responsive Reading — Psalm 23'),
        ('Opening Hymn', '#324 — "How Great Thou Art"'),
        ('Prayer of Confession', 'Silent Prayer followed by Corporate Confession'),
        ('Scripture Reading', 'Matthew 5:1-12 (ESV)'),
        ('Sermon', '"The Beatitudes: Living the Blessed Life" — Rev. Thomas Andrews'),
        ('Offering', 'Plates will be passed during the offertory hymn'),
        ('Closing Hymn', '#512 — "Blessed Assurance"'),
        ('Benediction', 'Rev. Thomas Andrews'),
        ('Postlude', 'Organ Music'),
    ]

    for item, desc in service_items:
        para = doc.add_paragraph()
        run_bold = para.add_run(f'{item}: ')
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_normal = para.add_run(desc)
        run_normal.font.size = Pt(10)

    # Page break to go to page 2
    doc.add_page_break()

    # --- Page 2: Announcements ---
    ann_title = doc.add_paragraph()
    ann_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ann_run = ann_title.add_run('COMMUNITY ANNOUNCEMENTS')
    ann_run.bold = True
    ann_run.font.size = Pt(14)

    doc.add_paragraph()  # Spacer

    # Food Pantry
    heading1 = doc.add_paragraph()
    h1_run = heading1.add_run('Food Pantry Drive')
    h1_run.bold = True
    h1_run.font.size = Pt(11)

    doc.add_paragraph(
        'Our annual spring food pantry drive is underway! We are collecting non-perishable '
        'food items to stock our community food bank. Drop-off bins are located in the '
        'narthex and fellowship hall. Most needed items include canned vegetables, pasta, '
        'rice, peanut butter, and canned soups. The drive runs through March 31st.'
    )

    # Youth Group
    heading2 = doc.add_paragraph()
    h2_run = heading2.add_run('Youth Group Events')
    h2_run.bold = True
    h2_run.font.size = Pt(11)

    doc.add_paragraph(
        'The St. Mark\'s Youth Group will hold a car wash fundraiser on Saturday, '
        'March 15th from 9:00 AM to 2:00 PM in the church parking lot. All proceeds '
        'go toward the summer mission trip to Appalachia. Contact Sarah Mitchell at '
        'smitchell@stmarks.org for more information or to volunteer.'
    )

    # Bible Study
    heading3 = doc.add_paragraph()
    h3_run = heading3.add_run('Bible Study — Wednesday Evenings')
    h3_run.bold = True
    h3_run.font.size = Pt(11)

    doc.add_paragraph(
        'Join Pastor Andrews for our Wednesday evening Bible study at 7:00 PM in Room 204. '
        'This month we are exploring the Book of Acts and the growth of the early church. '
        'All are welcome. Study guides are available at the Welcome Desk.'
    )

    # Prayer Chain
    heading4 = doc.add_paragraph()
    h4_run = heading4.add_run('Prayer Chain')
    h4_run.bold = True
    h4_run.font.size = Pt(11)

    doc.add_paragraph(
        'To add a prayer request to our chain, contact Deacon Margaret Holloway at '
        '(555) 423-7890 or email prayer@stmarks.org. All requests are kept strictly '
        'confidential. Our prayer team meets each Monday morning at 8:30 AM.'
    )

    # Upcoming Events
    heading5 = doc.add_paragraph()
    h5_run = heading5.add_run('Upcoming Events')
    h5_run.bold = True
    h5_run.font.size = Pt(11)

    events = [
        ('March 12', 'Wednesday — Bible Study, 7:00 PM, Room 204'),
        ('March 15', 'Saturday — Youth Car Wash, 9:00 AM–2:00 PM'),
        ('March 16', 'Sunday — Palm Sunday Celebration Service'),
        ('March 20', 'Thursday — Holy Thursday Service, 7:00 PM'),
        ('March 21', 'Good Friday Service, Noon and 7:00 PM'),
        ('March 23', 'Easter Sunday — Sunrise Service 6:30 AM; Main Service 10:00 AM'),
        ('March 31', 'Last day of Food Pantry Drive'),
    ]

    for date, event in events:
        para = doc.add_paragraph()
        run_date = para.add_run(f'{date}: ')
        run_date.bold = True
        run_date.font.size = Pt(10)
        run_event = para.add_run(event)
        run_event.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
