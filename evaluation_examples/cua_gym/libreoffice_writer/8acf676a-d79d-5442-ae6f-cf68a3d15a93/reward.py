"""
Reward Script: Create a label sheet with 30 labels (Avery 5160 layout).
Task ID: writer_lec_048
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with 10 rows x 3 cols (30 labels)
  Component 2 (0.25): Cell(0,0) contains 'SAMPLE - DO NOT MAIL'
  Component 3 (0.20): Cell(0,0) text is bold red
  Component 4 (0.30): Labels 2-30 contain company address in black text
"""

import os
from docx import Document
from docx.shared import RGBColor

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_048'

EXPECTED_ADDRESS = "GlobalTech LLC\n300 Commerce Blvd\nSuite 100\nAtlanta, GA 30301"
EXPECTED_LABEL1 = "SAMPLE - DO NOT MAIL"


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document — label sheet not created")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table has 10 rows x 3 cols (30 labels) (0.25 points)
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        total_cells = num_rows * num_cols
        if num_rows == 10 and num_cols == 3:
            print(f"PASS: Component 1 — Table is 10x3 = 30 labels (0.25 pts)")
            total_score += 0.25
        elif total_cells == 30:
            # Accept alternative layouts that still give 30 labels
            print(f"PASS: Component 1 — Table has 30 cells ({num_rows}x{num_cols}) (0.25 pts)")
            total_score += 0.25
        elif total_cells >= 28:
            # Close enough for partial credit
            print(f"PARTIAL: Component 1 — Table has {total_cells} cells ({num_rows}x{num_cols}), expected 30 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Table has {total_cells} cells ({num_rows}x{num_cols}), expected 30")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Cell(0,0) contains 'SAMPLE - DO NOT MAIL' (0.25 points)
    try:
        cell_00_text = table.cell(0, 0).text.strip()
        if EXPECTED_LABEL1 in cell_00_text:
            print(f"PASS: Component 2 — Cell(0,0) contains '{EXPECTED_LABEL1}' (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Cell(0,0) text is '{cell_00_text[:60]}', expected '{EXPECTED_LABEL1}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Cell(0,0) text is bold and red (0.20 points)
    try:
        cell_00 = table.cell(0, 0)
        # Collect formatting from all runs with text
        bold_runs = [run for para in cell_00.paragraphs for run in para.runs
                     if run.text.strip() and run.font.bold]
        red_runs = [run for para in cell_00.paragraphs for run in para.runs
                    if run.text.strip() and run.font.color and run.font.color.rgb
                    and isinstance(run.font.color.rgb, RGBColor)
                    and run.font.color.rgb[0] > 200
                    and run.font.color.rgb[1] < 100
                    and run.font.color.rgb[2] < 100]

        if len(bold_runs) > 0 and len(red_runs) > 0:
            print(f"PASS: Component 3 — Cell(0,0) text is bold and red (0.20 pts)")
            total_score += 0.20
        elif len(bold_runs) > 0:
            print(f"PARTIAL: Component 3 — Cell(0,0) is bold but not red (0.10 pts)")
            total_score += 0.10
        elif len(red_runs) > 0:
            print(f"PARTIAL: Component 3 — Cell(0,0) is red but not bold (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — Cell(0,0) text is not bold and not red")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Labels 2-30 contain company address in black text (0.30 points)
    try:
        address_count = 0
        black_count = 0
        total_address_cells = 0

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                # Skip the first cell (0,0) — that's label 1
                if ri == 0 and ci == 0:
                    continue
                total_address_cells += 1
                cell_text = cell.text.strip()

                # Normalize whitespace for comparison
                normalized_cell = "\n".join(line.strip() for line in cell_text.split("\n") if line.strip())
                normalized_expected = "\n".join(line.strip() for line in EXPECTED_ADDRESS.split("\n") if line.strip())

                if normalized_expected in normalized_cell or normalized_cell == normalized_expected:
                    address_count += 1

                    # Check text color is black (or default) — count non-black runs
                    non_black_runs = [
                        run for para in cell.paragraphs for run in para.runs
                        if run.text.strip() and run.font.color and run.font.color.rgb
                        and isinstance(run.font.color.rgb, RGBColor)
                        and not (run.font.color.rgb[0] < 50 and run.font.color.rgb[1] < 50 and run.font.color.rgb[2] < 50)
                    ]
                    if len(non_black_runs) == 0:
                        black_count += 1

        # Score based on proportion of correct address cells
        if total_address_cells > 0:
            address_ratio = address_count / 29  # expecting 29 address labels
            black_ratio = black_count / 29

            if address_count >= 28 and black_count >= 28:
                print(f"PASS: Component 4 — {address_count}/29 labels have correct address, {black_count}/29 in black (0.30 pts)")
                total_score += 0.30
            elif address_count >= 20:
                partial = 0.30 * (address_count / 29)
                print(f"PARTIAL: Component 4 — {address_count}/29 labels have correct address ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Only {address_count}/29 labels have correct address")
        else:
            print(f"FAIL: Component 4 — No address cells found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook — save any unsaved LibreOffice state
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Main execution
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
