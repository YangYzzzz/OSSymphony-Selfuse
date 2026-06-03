"""
Reward Script: Configure TOC with CHAPTER prefix and indentation levels
Task ID: writer_acad_058
Domain: libreoffice_writer
Scoring:
  Component 1: Level 1 entries have 'CHAPTER' prefix (0.40 points)
  Component 2: Level 2 entries have ~1 cm left indent (0.25 points)
  Component 3: Level 3 entries have ~2 cm left indent (0.25 points)
  Component 4: Tab stops with dot leaders preserved on all TOC entries (0.10 points)
"""

import os
from docx import Document
from docx.shared import Pt, Emu

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_058'

# Known TOC paragraph indices (P[15] through P[31])
TOC_START = 15
TOC_END = 31  # inclusive

# Level 1 entries (font size 152400 EMU = 12pt in initial, known chapter headings)
LEVEL1_TITLES = ['Introduction', 'Literature Review', 'Methodology', 'Results', 'Discussion', 'Conclusion']

# Level 2 entries (sub-chapter headings)
LEVEL2_TITLES = ['Research Design', 'Data Collection', 'Data Analysis',
                 'Quantitative Findings', 'Qualitative Findings',
                 'Implications for Theory', 'Practical Implications']

# Level 3 entries (sub-sub-chapter headings)
LEVEL3_TITLES = ['Survey Instruments', 'Interview Protocols',
                 'Descriptive Statistics', 'Regression Analysis']

# 1 cm in EMU = 360000; allow tolerance of ~10%
ONE_CM_EMU = 360000
TWO_CM_EMU = 720000
INDENT_TOLERANCE = 40000  # ~1mm tolerance


def classify_toc_entry(para_text):
    """Classify a TOC paragraph as level 1, 2, or 3 based on known titles."""
    # Strip any 'CHAPTER ' prefix for matching
    clean = para_text.replace('CHAPTER ', '').split('\t')[0].strip()
    if clean in LEVEL1_TITLES:
        return 1
    elif clean in LEVEL2_TITLES:
        return 2
    elif clean in LEVEL3_TITLES:
        return 3
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect TOC paragraphs
    toc_paras = []
    for i in range(TOC_START, min(TOC_END + 1, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.text.strip():  # skip empty
            toc_paras.append((i, para))

    if len(toc_paras) == 0:
        print("FAIL: No TOC entries found in expected paragraph range")
        print("REWARD: 0.0")
        return 0.0

    # Classify each TOC entry
    level1_paras = []
    level2_paras = []
    level3_paras = []
    for idx, para in toc_paras:
        level = classify_toc_entry(para.text)
        if level == 1:
            level1_paras.append((idx, para))
        elif level == 2:
            level2_paras.append((idx, para))
        elif level == 3:
            level3_paras.append((idx, para))

    # Component 1: Level 1 entries have 'CHAPTER' prefix (0.40 points)
    # This is the primary task-introduced change for Level 1 entries
    try:
        chapter_count = 0
        total_l1 = len(level1_paras)
        for idx, para in level1_paras:
            text = para.text.strip()
            if text.upper().startswith('CHAPTER ') or text.startswith('CHAPTER '):
                chapter_count += 1
                print(f"  PASS: P[{idx}] has CHAPTER prefix: '{text[:40]}'")
            else:
                print(f"  FAIL: P[{idx}] missing CHAPTER prefix: '{text[:40]}'")

        if total_l1 > 0 and chapter_count == total_l1:
            print(f"PASS: Component 1 -- All {total_l1} Level 1 entries have CHAPTER prefix (0.40 pts)")
            total_score += 0.40
        elif total_l1 > 0 and chapter_count > 0:
            partial = 0.40 * (chapter_count / total_l1)
            print(f"PARTIAL: Component 1 -- {chapter_count}/{total_l1} Level 1 entries have CHAPTER prefix ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- No Level 1 entries have CHAPTER prefix (0/{total_l1})")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Level 2 entries have ~1 cm left indent (0.25 points)
    # In initial env, left_indent is None. In golden, it should be ~360000 EMU (1 cm)
    try:
        indent_ok = 0
        total_l2 = len(level2_paras)
        for idx, para in level2_paras:
            left_indent = para.paragraph_format.left_indent
            if left_indent is not None:
                indent_val = int(left_indent)
                if abs(indent_val - ONE_CM_EMU) <= INDENT_TOLERANCE:
                    indent_ok += 1
                    print(f"  PASS: P[{idx}] Level 2 indent={indent_val} EMU (~1cm)")
                else:
                    print(f"  FAIL: P[{idx}] Level 2 indent={indent_val} EMU (expected ~{ONE_CM_EMU})")
            else:
                print(f"  FAIL: P[{idx}] Level 2 indent=None (expected ~{ONE_CM_EMU})")

        if total_l2 > 0 and indent_ok == total_l2:
            print(f"PASS: Component 2 -- All {total_l2} Level 2 entries have ~1cm indent (0.25 pts)")
            total_score += 0.25
        elif total_l2 > 0 and indent_ok > 0:
            partial = 0.25 * (indent_ok / total_l2)
            print(f"PARTIAL: Component 2 -- {indent_ok}/{total_l2} Level 2 entries have ~1cm indent ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- No Level 2 entries have ~1cm indent (0/{total_l2})")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Level 3 entries have ~2 cm left indent (0.25 points)
    # In initial env, left_indent is None. In golden, it should be ~720000 EMU (2 cm)
    try:
        indent_ok = 0
        total_l3 = len(level3_paras)
        for idx, para in level3_paras:
            left_indent = para.paragraph_format.left_indent
            if left_indent is not None:
                indent_val = int(left_indent)
                if abs(indent_val - TWO_CM_EMU) <= INDENT_TOLERANCE:
                    indent_ok += 1
                    print(f"  PASS: P[{idx}] Level 3 indent={indent_val} EMU (~2cm)")
                else:
                    print(f"  FAIL: P[{idx}] Level 3 indent={indent_val} EMU (expected ~{TWO_CM_EMU})")
            else:
                print(f"  FAIL: P[{idx}] Level 3 indent=None (expected ~{TWO_CM_EMU})")

        if total_l3 > 0 and indent_ok == total_l3:
            print(f"PASS: Component 3 -- All {total_l3} Level 3 entries have ~2cm indent (0.25 pts)")
            total_score += 0.25
        elif total_l3 > 0 and indent_ok > 0:
            partial = 0.25 * (indent_ok / total_l3)
            print(f"PARTIAL: Component 3 -- {indent_ok}/{total_l3} Level 3 entries have ~2cm indent ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 -- No Level 3 entries have ~2cm indent (0/{total_l3})")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Tab stops with dot leaders preserved on all TOC entries (0.10 points)
    # This is a preservation check -- it should pass on BOTH envs, so we use it as a gate:
    # Only award points if at least one task-change component (1, 2, or 3) also scored > 0
    # This ensures initial_env scores 0.0 overall (since components 1-3 fail on initial)
    try:
        dot_leader_count = 0
        total_toc = len(toc_paras)
        for idx, para in toc_paras:
            has_dot_leader = False
            for ts in para.paragraph_format.tab_stops:
                # Check for RIGHT-aligned tab with DOTS leader
                if str(ts.alignment) == 'RIGHT (2)' and str(ts.leader) == 'DOTS (1)':
                    has_dot_leader = True
                    break
            if has_dot_leader:
                dot_leader_count += 1

        # Gate: only count this if task-change score > 0
        task_change_score = total_score  # score from components 1-3
        if task_change_score > 0 and dot_leader_count == total_toc:
            print(f"PASS: Component 4 -- All {total_toc} TOC entries have right-aligned dot leaders (0.10 pts)")
            total_score += 0.10
        elif task_change_score > 0 and dot_leader_count > 0:
            partial = 0.10 * (dot_leader_count / total_toc)
            print(f"PARTIAL: Component 4 -- {dot_leader_count}/{total_toc} entries have dot leaders ({partial:.2f} pts)")
            total_score += partial
        elif task_change_score == 0:
            print(f"SKIP: Component 4 -- Gated on task-change score (currently 0.0)")
        else:
            print(f"FAIL: Component 4 -- No TOC entries have dot leaders")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
