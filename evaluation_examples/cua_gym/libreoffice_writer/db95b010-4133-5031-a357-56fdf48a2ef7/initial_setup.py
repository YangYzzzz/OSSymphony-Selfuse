"""
Initial Setup: Configure TOC in thesis document
Task ID: writer_acad_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_058'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_field(doc):
    """Add a TOC field code to the document."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.bold = True
    run.font.size = Pt(16)
    run.text = "TABLE OF CONTENTS"

    # Add empty paragraph as spacer
    doc.add_paragraph()

    # Add the TOC field code
    toc_para = doc.add_paragraph()
    run_begin = toc_para.add_run()
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run_begin._element.append(fld_char_begin)

    run_instr = toc_para.add_run()
    instr_text = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run_instr._element.append(instr_text)

    run_separate = toc_para.add_run()
    fld_char_separate = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
    )
    run_separate._element.append(fld_char_separate)

    # Add placeholder TOC entries (will be updated by LibreOffice)
    toc_entries = [
        ("1", "Introduction", "1"),
        ("1", "Literature Review", "5"),
        ("1", "Methodology", "12"),
        ("2", "Research Design", "12"),
        ("2", "Data Collection", "15"),
        ("3", "Survey Instruments", "15"),
        ("3", "Interview Protocols", "17"),
        ("2", "Data Analysis", "19"),
        ("1", "Results", "22"),
        ("2", "Quantitative Findings", "22"),
        ("3", "Descriptive Statistics", "22"),
        ("3", "Regression Analysis", "25"),
        ("2", "Qualitative Findings", "28"),
        ("1", "Discussion", "32"),
        ("2", "Implications for Theory", "32"),
        ("2", "Practical Implications", "35"),
        ("1", "Conclusion", "38"),
    ]

    for level, title, page in toc_entries:
        entry_para = doc.add_paragraph()
        entry_para.style = doc.styles[f'TOC Heading'] if False else None
        # Use default TOC style - no custom indentation
        pf = entry_para.paragraph_format
        # Default: minimal indent based on level (LibreOffice defaults)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        entry_run = entry_para.add_run(f"{title}")
        entry_run.font.size = Pt(11)
        if level == "1":
            entry_run.font.size = Pt(12)

        # Add tab + page number with dot leader
        tab_stops = pf.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        page_run = entry_para.add_run(f"\t{page}")
        page_run.font.size = Pt(11) if level != "1" else Pt(12)

    # End the TOC field
    run_end = toc_para.add_run()
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run_end._element.append(fld_char_end)

    # Page break after TOC
    doc.add_page_break()


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Title Page ---
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("The Impact of Digital Transformation on\nOrganizational Learning in Higher Education")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Times New Roman'

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_run = author_para.add_run("by\nElena Vasquez Rodriguez")
    author_run.font.size = Pt(14)
    author_run.font.name = 'Times New Roman'

    doc.add_paragraph()

    dept_para = doc.add_paragraph()
    dept_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_run = dept_para.add_run("A thesis submitted in partial fulfillment\nof the requirements for the degree of\nDoctor of Philosophy\n\nDepartment of Education Policy and Leadership\nUniversity of Western Ontario\n\nDecember 2025")
    dept_run.font.size = Pt(12)
    dept_run.font.name = 'Times New Roman'

    doc.add_page_break()

    # --- Table of Contents ---
    add_toc_field(doc)

    # --- Chapter 1: Introduction ---
    h1 = doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "The rapid advancement of digital technologies has fundamentally reshaped the landscape of "
        "higher education worldwide. Universities and colleges are increasingly leveraging digital "
        "tools, platforms, and strategies to enhance teaching, research, and administrative processes. "
        "This transformation extends beyond mere technological adoption; it represents a paradigm "
        "shift in how educational institutions create, share, and apply knowledge."
    )
    doc.add_paragraph(
        "Organizational learning, defined as the process through which institutions develop, retain, "
        "and transfer knowledge, has emerged as a critical factor in understanding how universities "
        "adapt to digital disruption. Despite growing interest in both digital transformation and "
        "organizational learning, relatively few studies have examined their intersection in the "
        "context of higher education."
    )

    # --- Chapter 2: Literature Review ---
    doc.add_page_break()
    h2 = doc.add_heading('Literature Review', level=1)
    doc.add_paragraph(
        "This chapter reviews the existing literature on digital transformation in higher education "
        "and organizational learning theory. The review is organized around three thematic areas: "
        "digital transformation frameworks, organizational learning models, and the intersection "
        "of these two domains in educational settings."
    )

    # --- Chapter 3: Methodology ---
    doc.add_page_break()
    h3 = doc.add_heading('Methodology', level=1)

    h3_1 = doc.add_heading('Research Design', level=2)
    doc.add_paragraph(
        "This study employed a mixed-methods research design combining quantitative survey data "
        "with qualitative interviews and document analysis. The convergent parallel design allowed "
        "for simultaneous collection of both data types, enabling comprehensive understanding of "
        "the phenomenon under investigation."
    )

    h3_2 = doc.add_heading('Data Collection', level=2)
    doc.add_paragraph(
        "Data were collected from twelve research-intensive universities across North America "
        "between September 2024 and March 2025. The sampling strategy employed purposive selection "
        "to ensure representation across institutional types, geographic regions, and stages of "
        "digital maturity."
    )

    h3_2_1 = doc.add_heading('Survey Instruments', level=3)
    doc.add_paragraph(
        "The quantitative component utilized a validated survey instrument adapted from the "
        "Digital Transformation Readiness Scale (DTRS; Chen & Williams, 2022) and the "
        "Organizational Learning Questionnaire (OLQ; Marsick & Watkins, 2003). The combined "
        "instrument contained 67 items measured on a 7-point Likert scale."
    )

    h3_2_2 = doc.add_heading('Interview Protocols', level=3)
    doc.add_paragraph(
        "Semi-structured interviews were conducted with 36 participants including provosts, "
        "chief information officers, faculty senate chairs, and instructional designers. Each "
        "interview lasted approximately 60-90 minutes and was audio-recorded with participant "
        "consent."
    )

    h3_3 = doc.add_heading('Data Analysis', level=2)
    doc.add_paragraph(
        "Quantitative data were analyzed using structural equation modeling (SEM) with AMOS 28.0. "
        "Qualitative data underwent thematic analysis following Braun and Clarke's (2006) six-phase "
        "framework. Integration of findings occurred through a joint display matrix approach."
    )

    # --- Chapter 4: Results ---
    doc.add_page_break()
    h4 = doc.add_heading('Results', level=1)

    h4_1 = doc.add_heading('Quantitative Findings', level=2)
    doc.add_paragraph(
        "The structural equation model demonstrated good fit indices (CFI = 0.94, RMSEA = 0.05, "
        "SRMR = 0.06), supporting the hypothesized relationships between digital transformation "
        "dimensions and organizational learning outcomes."
    )

    h4_1_1 = doc.add_heading('Descriptive Statistics', level=3)
    doc.add_paragraph(
        "A total of 1,247 valid responses were received, yielding a response rate of 34.2%. "
        "Respondents represented all twelve participating institutions, with faculty comprising "
        "52% of the sample, administrators 28%, and support staff 20%."
    )

    h4_1_2 = doc.add_heading('Regression Analysis', level=3)
    doc.add_paragraph(
        "Multiple regression analysis revealed that digital infrastructure investment (beta = 0.31, "
        "p < .001), leadership commitment (beta = 0.27, p < .001), and professional development "
        "opportunities (beta = 0.22, p < .01) were significant predictors of organizational "
        "learning capacity."
    )

    h4_2 = doc.add_heading('Qualitative Findings', level=2)
    doc.add_paragraph(
        "Thematic analysis of interview transcripts yielded five major themes: technological "
        "readiness, cultural resistance, distributed leadership, knowledge sharing networks, "
        "and adaptive capacity. Each theme is discussed in detail in the following subsections."
    )

    # --- Chapter 5: Discussion ---
    doc.add_page_break()
    h5 = doc.add_heading('Discussion', level=1)

    h5_1 = doc.add_heading('Implications for Theory', level=2)
    doc.add_paragraph(
        "The findings extend current understanding of organizational learning theory by "
        "demonstrating how digital transformation serves as both a catalyst and moderator of "
        "learning processes in higher education. The proposed Digital-Organizational Learning "
        "Framework (DOLF) integrates Senge's (1990) five disciplines with Vial's (2019) digital "
        "transformation framework."
    )

    h5_2 = doc.add_heading('Practical Implications', level=2)
    doc.add_paragraph(
        "University administrators should prioritize investment in digital infrastructure while "
        "simultaneously fostering cultures of experimentation and knowledge sharing. The study "
        "suggests that successful digital transformation requires alignment between technological "
        "capabilities, organizational culture, and strategic vision."
    )

    # --- Chapter 6: Conclusion ---
    doc.add_page_break()
    h6 = doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        "This study examined the relationship between digital transformation and organizational "
        "learning in higher education through a mixed-methods investigation of twelve research-"
        "intensive universities. The findings reveal that digital transformation positively "
        "influences organizational learning when supported by committed leadership, adequate "
        "infrastructure, and a culture that values continuous improvement."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
