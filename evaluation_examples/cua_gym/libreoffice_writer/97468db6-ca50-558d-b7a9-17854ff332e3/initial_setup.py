"""
Initial Setup: Read image editing checklist and edit scenery.jpg in GIMP
Task ID: osworld_multi_apps_writer_gimp_059
Domain: libreoffice_writer + gimp (multi-app)
"""

import os
import shlex
import subprocess
import time

# python-docx for .docx creation
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Pillow for scenery.jpg creation
from PIL import Image, ImageDraw, ImageFilter
import numpy as np

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_writer_gimp_059'
DESKTOP = '/home/user/Desktop'
DOC_OUTPUT = f'{WORKDIR}/photo_tasks.docx'
IMAGE_OUTPUT = f'{DESKTOP}/scenery.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_checklist_docx():
    """Create photo_tasks.docx with the image editing checklist."""
    doc = Document()

    # Title
    title = doc.add_heading("Photo Editing Checklist", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        "Please complete the following image editing tasks on 'scenery.jpg' "
        "using GIMP. Apply each step in order and export the final result."
    )
    intro.paragraph_format.space_after = Pt(12)

    # Subheading
    doc.add_heading("Editing Tasks", level=2)

    # Checklist items as numbered list
    items = [
        "Correct perspective distortion (rotate -2 degrees)",
        "Boost saturation by 30%",
        "Add a vignette effect",
        "Export as JPEG quality 90",
    ]

    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{item}")
        run.font.size = Pt(12)

    doc.add_paragraph()  # spacing

    # Export instructions section
    doc.add_heading("Export Instructions", level=2)
    export_para = doc.add_paragraph(
        "After completing all edits, export the finished image as "
        "'scenery_done.jpg' to the Desktop. Ensure the export quality is set to 90."
    )
    export_para.paragraph_format.space_before = Pt(6)

    # Notes section
    doc.add_heading("Notes", level=2)
    notes = [
        "The original file 'scenery.jpg' is located on the Desktop.",
        "Apply all effects sequentially before exporting.",
        "Do not overwrite the original 'scenery.jpg' — save as 'scenery_done.jpg'.",
    ]
    for note in notes:
        p = doc.add_paragraph(note, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    doc.save(DOC_OUTPUT)
    print(f'Checklist document created: {DOC_OUTPUT}')


def create_scenery_jpg():
    """Create a realistic-looking landscape scenery image."""
    os.makedirs(DESKTOP, exist_ok=True)

    width, height = 1280, 720
    img = Image.new("RGB", (width, height))
    pixels = img.load()

    # Sky gradient (top): light blue to deeper blue
    for y in range(height // 2):
        r = int(135 + (y / (height // 2)) * 50)
        g = int(206 - (y / (height // 2)) * 30)
        b = int(235 - (y / (height // 2)) * 20)
        for x in range(width):
            pixels[x, y] = (r, g, b)

    # Mountains in the middle
    for y in range(height // 2, int(height * 0.65)):
        factor = (y - height // 2) / (int(height * 0.65) - height // 2)
        r = int(80 + factor * 40)
        g = int(100 + factor * 60)
        b = int(60 + factor * 30)
        for x in range(width):
            pixels[x, y] = (r, g, b)

    # Grass/meadow (lower portion)
    for y in range(int(height * 0.65), height):
        factor = (y - int(height * 0.65)) / (height - int(height * 0.65))
        r = int(34 + factor * 30)
        g = int(139 - factor * 20)
        b = int(34 - factor * 10)
        for x in range(width):
            pixels[x, y] = (r, g, b)

    draw = ImageDraw.Draw(img)

    # Draw mountain silhouettes
    mountain_points = [
        (0, int(height * 0.7)),
        (100, int(height * 0.45)),
        (200, int(height * 0.55)),
        (350, int(height * 0.35)),
        (500, int(height * 0.5)),
        (650, int(height * 0.3)),
        (800, int(height * 0.48)),
        (950, int(height * 0.4)),
        (1100, int(height * 0.52)),
        (1280, int(height * 0.45)),
        (1280, height),
        (0, height),
    ]
    draw.polygon(mountain_points, fill=(70, 90, 50))

    # Draw some clouds
    for cx, cy, cr in [(200, 80, 60), (500, 60, 80), (900, 90, 55), (1100, 70, 65)]:
        for dx, dy, dr in [(-cr//2, 0, cr), (0, -cr//3, int(cr*0.8)), (cr//2, 0, cr)]:
            draw.ellipse(
                [cx + dx - dr//2, cy + dy - dr//2, cx + dx + dr//2, cy + dy + dr//2],
                fill=(240, 248, 255)
            )

    # Draw a river/stream
    river_points = [(400, height), (380, int(height * 0.75)), (420, int(height * 0.65)),
                    (440, int(height * 0.6))]
    for i in range(len(river_points) - 1):
        draw.line([river_points[i], river_points[i+1]], fill=(135, 206, 235), width=8)

    # Add some trees (simple triangles)
    tree_positions = [(150, int(height * 0.6)), (250, int(height * 0.62)),
                      (700, int(height * 0.58)), (800, int(height * 0.6)),
                      (1050, int(height * 0.62))]
    for tx, ty in tree_positions:
        tree_h = 60
        tree_w = 25
        draw.polygon(
            [(tx, ty - tree_h), (tx - tree_w, ty), (tx + tree_w, ty)],
            fill=(0, 80, 0)
        )
        draw.rectangle([tx - 5, ty, tx + 5, ty + 20], fill=(80, 40, 20))

    # Apply a slight blur for realistic depth-of-field effect
    img = img.filter(ImageFilter.GaussianBlur(radius=0.5))

    # Save as JPEG quality 95 (high quality initial)
    img.save(IMAGE_OUTPUT, "JPEG", quality=95)
    print(f'Scenery image created: {IMAGE_OUTPUT}')


def create_initial():
    create_checklist_docx()
    create_scenery_jpg()

    # GUI-ready startup: open photo_tasks.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=2.0)
    # Also open GIMP with the scenery image so the agent can follow the checklist
    launch_gui(f'gimp "{IMAGE_OUTPUT}"', delay_sec=3.0)

    print('GUI_READY: launched LibreOffice Writer and GIMP with DISPLAY=:0')


create_initial()
