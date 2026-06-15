"""
Initial Setup: Create Client_Contacts.docx with a 5-column x 15-row contact table
Task ID: writer_tm_027
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Client Contact Directory", level=1)

    doc.add_paragraph(
        "Below is the master contact list for all active clients. "
        "Please keep this document updated with any address or phone changes."
    )

    # 5 columns x 15 rows (1 header + 14 data)
    table = doc.add_table(rows=15, cols=5)
    table.style = "Table Grid"

    # Headers
    headers = ['Name', 'Phone', 'Fax Number', 'Email', 'Address']
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # 14 rows of realistic client data
    data = [
        ['Elena Vasquez', '(415) 555-0142', '(415) 555-0143', 'elena.vasquez@meridianlaw.com', '1200 Market St, Suite 400, San Francisco, CA 94103'],
        ['Robert Nakamura', '(212) 555-0178', '(212) 555-0179', 'r.nakamura@techbridge.io', '350 Fifth Ave, Floor 22, New York, NY 10118'],
        ['Sarah Chen', '(312) 555-0231', '(312) 555-0232', 'schen@windycityconsulting.com', '233 S Wacker Dr, Chicago, IL 60606'],
        ['Marcus Thompson', '(617) 555-0394', '(617) 555-0395', 'marcus.t@bostonglobal.net', '75 State St, Suite 1100, Boston, MA 02109'],
        ['Priya Sharma', '(650) 555-0456', '(650) 555-0457', 'priya@valleyinnovations.com', '2580 Sand Hill Rd, Menlo Park, CA 94025'],
        ['David Okonkwo', '(404) 555-0512', '(404) 555-0513', 'd.okonkwo@peachtreeadvisors.com', '191 Peachtree St NE, Atlanta, GA 30303'],
        ['Jennifer Liu', '(206) 555-0623', '(206) 555-0624', 'jliu@pacificrimtrading.com', '701 Pike St, Suite 900, Seattle, WA 98101'],
        ['Carlos Rivera', '(305) 555-0789', '(305) 555-0790', 'carlos@suncoastlogistics.com', '100 SE 2nd St, Miami, FL 33131'],
        ['Amanda Foster', '(720) 555-0845', '(720) 555-0846', 'afoster@rockymountainhr.com', '1600 Broadway, Suite 2100, Denver, CO 80202'],
        ['Wei Zhang', '(469) 555-0901', '(469) 555-0902', 'wzhang@lonestarfinancial.com', '2200 Ross Ave, Dallas, TX 75201'],
        ['Rachel Goldstein', '(215) 555-1034', '(215) 555-1035', 'rgoldstein@libertymedical.org', '1500 Market St, Philadelphia, PA 19102'],
        ['Tomasz Kowalski', '(602) 555-1167', '(602) 555-1168', 'tkowalski@desertedge.com', '2 N Central Ave, Phoenix, AZ 85004'],
        ['Monica Diaz', '(503) 555-1289', '(503) 555-1290', 'mdiaz@cascadedesignstudio.com', '111 SW 5th Ave, Portland, OR 97204'],
        ['James O\'Brien', '(612) 555-1345', '(612) 555-1346', 'jobrien@northlandpartners.com', '225 S 6th St, Suite 3900, Minneapolis, MN 55402'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx, col_idx).text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
