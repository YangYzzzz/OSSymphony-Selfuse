"""
Reward Script: Delete column 3 ('Fax Number') from client contact table
Task ID: writer_tm_027
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table has exactly 4 columns (reduced from 5)
  Component 2 (0.3): 'Fax Number' header is absent
  Component 3 (0.2): Remaining headers are ['Name', 'Phone', 'Email', 'Address'] in order
  Component 4 (0.2): Data in remaining columns is preserved (spot-check rows)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_027'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table has exactly 4 columns (0.3 points)
    # Initial has 5 columns; golden should have 4 after removing 'Fax Number'
    try:
        if num_cols == 4:
            print(f"PASS: Component 1 — Table has 4 columns (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Expected 4 columns, found {num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'Fax Number' header is absent (0.3 points)
    # The header row should not contain 'Fax Number' anywhere
    try:
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        fax_present = any('fax' in h.lower() for h in header_cells)
        if not fax_present:
            print(f"PASS: Component 2 — 'Fax Number' not in headers: {header_cells} (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — 'Fax Number' still present in headers: {header_cells}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Remaining headers are ['Name', 'Phone', 'Email', 'Address'] in order (0.2 points)
    try:
        expected_headers = ['Name', 'Phone', 'Email', 'Address']
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        if header_cells == expected_headers:
            print(f"PASS: Component 3 — Headers match expected order: {header_cells} (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Expected {expected_headers}, found {header_cells}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Data preserved in 4-column table (0.2 points)
    # Compound check: table must have 4 columns AND data in remaining columns is correct.
    # This ensures it only passes when the fax column was removed (4 cols) AND data survived.
    try:
        if num_cols != 4:
            print(f"FAIL: Component 4 — Table does not have 4 columns ({num_cols}), skipping data check")
        else:
            spot_checks = {
                1: {0: 'Elena Vasquez', 1: '(415) 555-0142', 2: 'elena.vasquez@meridianlaw.com', 3: '1200 Market St, Suite 400, San Francisco, CA 94103'},
                7: {0: 'Jennifer Liu', 1: '(206) 555-0623', 2: 'jliu@pacificrimtrading.com', 3: '701 Pike St, Suite 900, Seattle, WA 98101'},
                14: {0: "James O'Brien", 1: '(612) 555-1345', 2: 'jobrien@northlandpartners.com', 3: '225 S 6th St, Suite 3900, Minneapolis, MN 55402'},
            }
            checks_passed = 0
            checks_total = 0
            for row_idx, col_vals in spot_checks.items():
                if row_idx >= num_rows:
                    print(f"WARN: Row {row_idx} does not exist (only {num_rows} rows)")
                    continue
                row_cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
                for ci, expected_val in col_vals.items():
                    checks_total += 1
                    if ci < len(row_cells) and row_cells[ci] == expected_val:
                        checks_passed += 1
                    else:
                        actual = row_cells[ci] if ci < len(row_cells) else 'N/A'
                        print(f"  DATA MISMATCH: Row {row_idx} Col {ci}: expected '{expected_val}', got '{actual}'")

            if checks_total > 0 and checks_passed == checks_total:
                print(f"PASS: Component 4 — 4-col table with all {checks_passed}/{checks_total} data checks passed (0.2 pts)")
                total_score += 0.2
            elif checks_total > 0:
                partial = 0.2 * (checks_passed / checks_total)
                print(f"PARTIAL: Component 4 — {checks_passed}/{checks_total} data checks passed ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — No spot-checks could be performed")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
