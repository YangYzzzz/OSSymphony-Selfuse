"""
Initial Setup: Create a Writer document with an outline numbered list (all level 1).
Task ID: writer_lec_014
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_014'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_multilevel_numbering(doc):
    """
    Create a multi-level numbering definition in the document.
    Level 0: "1.", "2.", "3." ...
    Level 1: "1.1", "1.2", "2.1" ...
    """
    # Get or create numbering part
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part._element

    # Create abstractNum with multi-level definition
    abstractNum_xml = f'''
    <w:abstractNum w:abstractNumId="10"
        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:multiLevelType w:val="multilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="360" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                <w:sz w:val="24"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                <w:sz w:val="24"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="1080" w:hanging="360"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    '''
    abstractNum = parse_xml(abstractNum_xml)
    numbering_elem.append(abstractNum)

    # Create num referencing the abstractNum
    num_xml = '''
    <w:num w:numId="10"
        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:abstractNumId w:val="10"/>
    </w:num>
    '''
    num = parse_xml(num_xml)
    numbering_elem.append(num)

    return 10  # numId


def add_list_item(doc, text, num_id, level=0):
    """Add a numbered list item at the specified level."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    # Set numbering properties
    pPr = para._element.get_or_add_pPr()
    numPr = parse_xml(
        f'<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'  <w:ilvl w:val="{level}"/>'
        f'  <w:numId w:val="{num_id}"/>'
        f'</w:numPr>'
    )
    pPr.insert(0, numPr)

    return para


def create_initial():
    doc = Document()

    # Add a title
    title = doc.add_heading('Project Development Outline', level=1)

    # Add a brief intro paragraph
    intro = doc.add_paragraph(
        'The following outline describes the major phases of the software '
        'development project for the Q3 2025 product release cycle.'
    )
    intro_run = intro.runs[0]
    intro_run.font.name = 'Calibri'
    intro_run.font.size = Pt(11)

    # Add spacing before the list
    doc.add_paragraph()

    # Create multi-level numbering
    num_id = create_multilevel_numbering(doc)

    # All items at level 1 (ilvl=0) in the initial state
    items = [
        'Frontend Design',
        'Backend Development',
        'Database Setup',
        'API Configuration',
        'Testing',
    ]

    for item in items:
        add_list_item(doc, item, num_id, level=0)

    # Add some closing text
    doc.add_paragraph()
    closing = doc.add_paragraph(
        'Each phase should be completed sequentially with review checkpoints '
        'between major milestones. Team leads are responsible for coordinating '
        'dependencies across phases.'
    )
    closing_run = closing.runs[0]
    closing_run.font.name = 'Calibri'
    closing_run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
