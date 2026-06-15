"""
Reward Script: Remove horizontal inner borders from table
Task ID: writer_tm_041
Domain: libreoffice_writer
Scoring:
  Component 1 (0.50): insideH border is explicitly set to 'none'/'nil' in tblBorders
  Component 2 (0.25): insideH removed AND insideV remains 'single' sz=4
  Component 3 (0.25): insideH removed AND all 4 outer borders remain 'single' sz=4

All components are anchored to the insideH explicit removal.
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_041'


def persist_app_state(domain: str):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_border_info(tbl_borders, border_name):
    """Get a table border element by name, return (val, sz) or (None, None) if absent."""
    if tbl_borders is None:
        return (None, None)
    el = tbl_borders.find(qn(f'w:{border_name}'))
    if el is None:
        return (None, None)
    val = el.get(qn('w:val'))
    sz = el.get(qn('w:sz'))
    return (val, sz)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one table with 8 rows and 4 columns
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) != 8 or len(table.columns) != 4:
        print(f"FAIL: Table size is {len(table.rows)}x{len(table.columns)}, expected 8x4")
        print("REWARD: 0.0")
        return 0.0

    tblPr = table._tbl.find(qn('w:tblPr'))
    tblBorders = tblPr.find(qn('w:tblBorders')) if tblPr is not None else None

    # Strategy: The task requires EXPLICIT removal of insideH borders.
    # In the golden file, tblBorders exists with insideH val='none'.
    # In the initial file, tblBorders may be absent (borders inherited from TableGrid style)
    # or present with insideH val='single'.
    #
    # We check two approaches for detecting insideH removal:
    # A) Table-level: tblBorders exists AND insideH val in ('none', 'nil')
    # B) Cell-level: majority of inner cells have bottom border set to none/nil

    insideH_removed = False
    removal_method = ""

    # Approach A: Table-level check
    if tblBorders is not None:
        insideH_val, insideH_sz = get_border_info(tblBorders, 'insideH')
        if insideH_val in ('none', 'nil'):
            insideH_removed = True
            removal_method = f"table-level insideH val='{insideH_val}'"
        elif insideH_val is None:
            # insideH element absent within tblBorders — check if it was removed
            insideH_el = tblBorders.find(qn('w:insideH'))
            if insideH_el is None:
                # Element absent but other borders present could mean it was removed
                # Check if other borders ARE present (confirming selective removal)
                other_borders_present = 0
                for name in ['top', 'left', 'bottom', 'right', 'insideV']:
                    el = tblBorders.find(qn(f'w:{name}'))
                    if el is not None and el.get(qn('w:val')) == 'single':
                        other_borders_present += 1
                if other_borders_present >= 3:
                    insideH_removed = True
                    removal_method = "table-level insideH element absent while other borders present"

    # Approach B: Cell-level check (fallback)
    if not insideH_removed:
        cells_with_none_h = 0
        total_inner_cells = 0
        for ri in range(len(table.rows) - 1):
            for ci in range(len(table.columns)):
                total_inner_cells += 1
                cell = table.cell(ri, ci)
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        bottom_el = tcBorders.find(qn('w:bottom'))
                        if bottom_el is not None:
                            bval = bottom_el.get(qn('w:val'))
                            if bval in ('none', 'nil'):
                                cells_with_none_h += 1
        if total_inner_cells > 0 and cells_with_none_h >= total_inner_cells * 0.8:
            insideH_removed = True
            removal_method = f"cell-level ({cells_with_none_h}/{total_inner_cells} cells have bottom=none)"

    # Component 1: insideH border is removed (0.50 points)
    try:
        if insideH_removed:
            print(f"PASS: Component 1 — horizontal inner borders removed ({removal_method}) (0.50 pts)")
            total_score += 0.50
        else:
            if tblBorders is not None:
                val, sz = get_border_info(tblBorders, 'insideH')
                print(f"FAIL: Component 1 — insideH val='{val}' sz='{sz}', expected 'none' or absent")
            else:
                print(f"FAIL: Component 1 — no tblBorders found; borders inherited from style (not removed)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: insideH removed AND insideV remains single/4 (0.25 points)
    # Anchored to insideH change
    try:
        if insideH_removed:
            insideV_val, insideV_sz = get_border_info(tblBorders, 'insideV')
            if insideV_val == 'single' and insideV_sz == '4':
                print(f"PASS: Component 2 — insideH removed AND insideV intact (single/4) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — insideH removed but insideV val='{insideV_val}' sz='{insideV_sz}', expected single/4")
        else:
            print(f"FAIL: Component 2 — insideH not removed, skipping insideV check")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: insideH removed AND all outer borders remain single/4 (0.25 points)
    # Anchored to insideH change
    try:
        if insideH_removed:
            outer_names = ['top', 'left', 'bottom', 'right']
            outer_ok = 0
            for name in outer_names:
                val, sz = get_border_info(tblBorders, name)
                if val == 'single' and sz == '4':
                    outer_ok += 1
                else:
                    print(f"  DETAIL: outer border '{name}' val='{val}' sz='{sz}' — not single/4")

            if outer_ok == 4:
                print(f"PASS: Component 3 — insideH removed AND all 4 outer borders intact (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — only {outer_ok}/4 outer borders are single/4")
        else:
            print(f"FAIL: Component 3 — insideH not removed, skipping outer border check")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification
persist_app_state("libreoffice_writer")

# Run verification
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
