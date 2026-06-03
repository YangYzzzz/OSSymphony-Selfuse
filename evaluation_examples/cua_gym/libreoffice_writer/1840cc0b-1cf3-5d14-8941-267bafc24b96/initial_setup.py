"""
Initial Setup: Create a document with a 4x8 table with all borders 0.5pt solid black.
Task ID: writer_tm_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell using XML manipulation.
    Each border param is a dict with keys: sz, val, color, space.
    sz is in eighths of a point (4 = 0.5pt).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)

    for border_name, border_attrs in [('top', top), ('bottom', bottom),
                                       ('left', left), ('right', right)]:
        if border_attrs is not None:
            element = tcBorders.find(qn(f'w:{border_name}'))
            if element is not None:
                tcBorders.remove(element)
            border_el = parse_xml(
                f'<w:{border_name} {nsdecls("w")} '
                f'w:val="{border_attrs["val"]}" '
                f'w:sz="{border_attrs["sz"]}" '
                f'w:space="{border_attrs.get("space", 0)}" '
                f'w:color="{border_attrs["color"]}"/>'
            )
            tcBorders.append(border_el)


def set_table_borders(table, sz=4, val="single", color="000000"):
    """Set uniform borders on the entire table via tblBorders XML."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    # Remove existing tblBorders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading("Q1 2025 Regional Sales Report", level=1)

    # Introductory paragraph
    doc.add_paragraph(
        "The following table summarizes quarterly sales performance by product "
        "category across our four main regions. All figures are in thousands of USD."
    )

    # Create 4-column x 8-row table
    table = doc.add_table(rows=8, cols=4)
    table.style = "Table Grid"

    # Define data: header + 7 data rows
    data = [
        ["Product Category", "North Region", "East Region", "West Region"],
        ["Consumer Electronics", "$145,200", "$132,800", "$158,400"],
        ["Office Supplies", "$67,350", "$72,100", "$59,800"],
        ["Industrial Equipment", "$234,500", "$198,700", "$221,300"],
        ["Software Licenses", "$89,400", "$95,200", "$102,600"],
        ["Maintenance Services", "$42,750", "$38,600", "$45,100"],
        ["Consulting Fees", "$78,900", "$82,300", "$71,500"],
        ["Training Programs", "$23,400", "$27,800", "$31,200"],
    ]

    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = val
            # Bold the header row
            if r == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(11)

    # Set all borders to 0.5pt solid black (sz=4 means 4 eighths of a point = 0.5pt)
    set_table_borders(table, sz=4, val="single", color="000000")

    # Add a closing paragraph
    doc.add_paragraph(
        "Note: Figures are preliminary and subject to final audit review."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
