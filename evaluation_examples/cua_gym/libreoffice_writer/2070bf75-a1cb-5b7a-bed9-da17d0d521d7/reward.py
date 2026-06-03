"""
Reward Script: Set up footer with 'Page X of Y' format centered at bottom
Task ID: writer_tech_041
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Footer contains 'Page' and 'of' text fragments
  Component 2 (0.3): Footer has PAGE field code for current page number
  Component 3 (0.2): Footer has NUMPAGES field code for total page count
  Component 4 (0.2): Footer paragraph is center-aligned
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_041'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section
    if len(doc.sections) == 0:
        print("FAIL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    footer = section.footer

    # Precondition: footer must have at least one paragraph
    if not footer.paragraphs:
        print("FAIL: Footer has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    footer_para = footer.paragraphs[0]
    footer_text = footer_para.text

    # Extract field codes from footer XML
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    instr_texts = footer_para._element.findall('.//w:instrText', ns)
    field_codes = [it.text.strip() for it in instr_texts]

    print(f"Footer text: {repr(footer_text)}")
    print(f"Footer field codes: {field_codes}")
    print(f"Footer alignment: {footer_para.paragraph_format.alignment}")

    # Component 1: Footer contains 'Page' and 'of' text (0.3 points)
    # This checks the static text fragments that form "Page X of Y"
    try:
        text_lower = footer_text.lower()
        has_page_text = 'page' in text_lower
        has_of_text = ' of ' in text_lower
        if has_page_text and has_of_text:
            print(f"PASS: Component 1 — Footer contains 'Page' and 'of' text (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Expected 'Page' and 'of' in footer, found: {repr(footer_text)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Footer has PAGE field code (0.3 points)
    # The PAGE field inserts the current page number dynamically
    try:
        has_page_field = any('PAGE' in fc and 'NUMPAGES' not in fc for fc in field_codes)
        if has_page_field:
            print(f"PASS: Component 2 — PAGE field code found in footer (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — No PAGE field code in footer. Fields found: {field_codes}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Footer has NUMPAGES field code (0.2 points)
    # The NUMPAGES field inserts the total page count dynamically
    try:
        has_numpages_field = any('NUMPAGES' in fc for fc in field_codes)
        if has_numpages_field:
            print(f"PASS: Component 3 — NUMPAGES field code found in footer (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — No NUMPAGES field code in footer. Fields found: {field_codes}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Footer paragraph is center-aligned (0.2 points)
    # Task requires the footer to be centered at the bottom of each page
    try:
        alignment = footer_para.paragraph_format.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            print(f"PASS: Component 4 — Footer is center-aligned (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 4 — Expected CENTER alignment, found: {alignment}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
