"""
Initial Setup: Create a 12-page technical document with no footer
Task ID: writer_tech_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a manual page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Make sure there is NO footer
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""

    # --- Page 1: Title Page ---
    doc.add_heading('Cloud Infrastructure Migration Guide', level=0)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Prepared by the Platform Engineering Team')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run('Version 3.2 | March 2026')
    run2.font.size = Pt(12)
    run2.font.italic = True

    p3 = doc.add_paragraph()
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p3.add_run('Confidential - Internal Use Only').font.size = Pt(10)

    add_page_break(doc)

    # --- Page 2: Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Current Architecture Assessment',
        '3. Migration Strategy Overview',
        '4. Phase 1: Infrastructure Preparation',
        '5. Phase 2: Data Migration',
        '6. Phase 3: Application Migration',
        '7. Phase 4: Testing and Validation',
        '8. Rollback Procedures',
        '9. Security Considerations',
        '10. Cost Analysis and Budget',
        '11. Timeline and Milestones',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    add_page_break(doc)

    # --- Page 3: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines the comprehensive migration strategy for transitioning '
        'our on-premises infrastructure to a hybrid cloud environment. The migration '
        'encompasses 47 production services, 12 databases, and approximately 8.5 TB of '
        'persistent storage currently hosted across three data centers in Singapore, '
        'Frankfurt, and Virginia.'
    )
    doc.add_paragraph(
        'The estimated timeline for complete migration is 18 months, divided into four '
        'distinct phases. Total projected cost savings over a 3-year period are estimated '
        'at $2.4 million, with an initial capital expenditure of $680,000 for tooling, '
        'training, and parallel infrastructure operation during the transition period.'
    )
    doc.add_paragraph(
        'Key stakeholders include the CTO office, DevOps engineering, security operations, '
        'and the data platform team. Executive sponsorship has been confirmed by VP of '
        'Engineering Sarah Chen and CTO David Rodriguez.'
    )

    add_page_break(doc)

    # --- Page 4: Current Architecture ---
    doc.add_heading('2. Current Architecture Assessment', level=1)
    doc.add_heading('2.1 Server Inventory', level=2)
    doc.add_paragraph(
        'Our current on-premises deployment consists of 124 physical servers distributed '
        'across three geographic regions. The Singapore data center hosts 52 servers '
        'primarily running customer-facing APIs and the real-time analytics pipeline. '
        'Frankfurt operates 38 servers handling European customer data in compliance with '
        'GDPR requirements. Virginia maintains 34 servers for North American services and '
        'disaster recovery.'
    )

    # Add a table for server inventory
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Data Center', 'Server Count', 'CPU Cores', 'RAM (TB)']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
    data = [
        ['Singapore (SG-1)', '52', '832', '4.2'],
        ['Frankfurt (EU-1)', '38', '608', '3.1'],
        ['Virginia (US-1)', '34', '544', '2.8'],
        ['Total', '124', '1,984', '10.1'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_heading('2.2 Network Topology', level=2)
    doc.add_paragraph(
        'Inter-datacenter connectivity is maintained through dedicated 10 Gbps links '
        'provided by Equinix and Megaport. Average latency between SG-1 and EU-1 is '
        '165ms, while SG-1 to US-1 averages 210ms. Internal datacenter networking uses '
        'a spine-leaf architecture with 25 Gbps uplinks to each top-of-rack switch.'
    )

    add_page_break(doc)

    # --- Page 5: Migration Strategy ---
    doc.add_heading('3. Migration Strategy Overview', level=1)
    doc.add_paragraph(
        'We will employ a hybrid migration strategy combining lift-and-shift for stateless '
        'services with re-architecture for data-intensive workloads. The "6 Rs" framework '
        'has been applied to categorize each service: Rehost (18 services), Replatform '
        '(14 services), Refactor (8 services), Repurchase (4 services), Retain (2 services), '
        'and Retire (1 service).'
    )
    doc.add_heading('3.1 Decision Criteria', level=2)
    criteria = [
        'Service complexity and inter-dependency mapping',
        'Data sovereignty and compliance requirements (GDPR, SOC2, PCI-DSS)',
        'Performance SLAs and latency sensitivity',
        'Database coupling and state management patterns',
        'Team readiness and cloud-native expertise assessment',
    ]
    for item in criteria:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Target Architecture', level=2)
    doc.add_paragraph(
        'The target state leverages Kubernetes (EKS) for container orchestration, '
        'Aurora PostgreSQL for relational databases, ElastiCache for session management, '
        'and S3 with CloudFront for static asset delivery. Service mesh will be implemented '
        'using Istio for inter-service communication and observability.'
    )

    add_page_break(doc)

    # --- Page 6: Phase 1 ---
    doc.add_heading('4. Phase 1: Infrastructure Preparation', level=1)
    doc.add_heading('4.1 Cloud Account Structure', level=2)
    doc.add_paragraph(
        'A multi-account strategy using AWS Organizations will be implemented with separate '
        'accounts for production, staging, development, security, and shared services. '
        'Service Control Policies (SCPs) will enforce guardrails across all accounts.'
    )
    doc.add_heading('4.2 Network Foundation', level=2)
    doc.add_paragraph(
        'Transit Gateway will connect VPCs across three regions (ap-southeast-1, eu-central-1, '
        'us-east-1) mirroring our current datacenter topology. Direct Connect circuits of '
        '10 Gbps will be established to each datacenter for hybrid connectivity during the '
        'transition period. VPC CIDR ranges have been planned to avoid conflicts with '
        'existing on-premises 10.0.0.0/8 allocations.'
    )
    doc.add_heading('4.3 Identity and Access Management', level=2)
    doc.add_paragraph(
        'AWS IAM Identity Center (SSO) will integrate with our existing Okta directory. '
        'Role-based access control will map to current LDAP groups. Service accounts will '
        'use IAM roles with IRSA (IAM Roles for Service Accounts) for Kubernetes pods.'
    )

    add_page_break(doc)

    # --- Page 7: Phase 2 ---
    doc.add_heading('5. Phase 2: Data Migration', level=1)
    doc.add_heading('5.1 Database Migration', level=2)
    doc.add_paragraph(
        'The 12 PostgreSQL databases totaling 3.2 TB will be migrated using AWS Database '
        'Migration Service (DMS) with change data capture (CDC) for zero-downtime migration. '
        'The customer_transactions database (1.8 TB) requires special handling due to its '
        'real-time replication requirements and complex trigger logic.'
    )

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    h2 = ['Database', 'Size (GB)', 'Priority', 'Strategy']
    for i, h in enumerate(h2):
        table2.cell(0, i).text = h
        for run in table2.cell(0, i).paragraphs[0].runs:
            run.bold = True
    db_data = [
        ['customer_transactions', '1,843', 'Critical', 'DMS with CDC'],
        ['product_catalog', '456', 'High', 'DMS batch'],
        ['user_profiles', '312', 'High', 'DMS with CDC'],
        ['analytics_warehouse', '398', 'Medium', 'S3 export/import'],
        ['audit_logs', '224', 'Low', 'Archive to S3'],
    ]
    for r, row_data in enumerate(db_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_heading('5.2 Object Storage Migration', level=2)
    doc.add_paragraph(
        'Approximately 5.3 TB of object storage (documents, images, and media files) will '
        'be migrated to S3 using AWS DataSync. Initial bulk transfer will use AWS Snowball '
        'Edge for the Singapore datacenter due to bandwidth constraints. Subsequent '
        'incremental syncs will occur over Direct Connect.'
    )

    add_page_break(doc)

    # --- Page 8: Phase 3 ---
    doc.add_heading('6. Phase 3: Application Migration', level=1)
    doc.add_heading('6.1 Containerization', level=2)
    doc.add_paragraph(
        'All 47 services will be containerized using multi-stage Docker builds. Base images '
        'will be maintained in ECR with automated vulnerability scanning via Trivy. '
        'The payment-gateway and fraud-detection services require special attention due to '
        'PCI-DSS compliance requirements for container hardening.'
    )
    doc.add_heading('6.2 Kubernetes Deployment', level=2)
    doc.add_paragraph(
        'EKS clusters will be provisioned using Terraform with the following configuration: '
        'managed node groups with c6i.2xlarge instances for compute workloads, r6i.xlarge '
        'for memory-intensive services, and Karpenter for auto-scaling. Helm charts will '
        'standardize deployment manifests across all services.'
    )
    doc.add_heading('6.3 Migration Wave Planning', level=2)
    doc.add_paragraph(
        'Services are grouped into five migration waves based on dependency analysis and '
        'risk assessment. Wave 1 includes 8 stateless API services with no database '
        'dependencies. Wave 5 includes the real-time analytics pipeline and the payment '
        'processing system, which have the most complex dependency graphs.'
    )

    add_page_break(doc)

    # --- Page 9: Phase 4 ---
    doc.add_heading('7. Phase 4: Testing and Validation', level=1)
    doc.add_heading('7.1 Performance Benchmarks', level=2)
    doc.add_paragraph(
        'Each migrated service must meet or exceed current performance baselines. Key '
        'metrics include p99 latency, throughput (requests/second), and error rates. '
        'Load testing will be conducted using k6 with production traffic replay captured '
        'via GoReplay. Acceptance criteria: p99 latency within 10% of baseline, zero '
        'increase in error rates, throughput capacity at 150% of peak observed traffic.'
    )
    doc.add_heading('7.2 Disaster Recovery Testing', level=2)
    doc.add_paragraph(
        'Full DR tests will be executed quarterly during the migration period. RTO target '
        'is 4 hours for tier-1 services and 8 hours for tier-2 services. RPO target is '
        '15 minutes for all transactional databases. Chaos engineering experiments using '
        'Gremlin will validate resilience of the new infrastructure.'
    )

    add_page_break(doc)

    # --- Page 10: Rollback Procedures ---
    doc.add_heading('8. Rollback Procedures', level=1)
    doc.add_paragraph(
        'Each migration wave includes a documented rollback plan with pre-defined triggers. '
        'Rollback triggers include: sustained error rate above 1% for more than 5 minutes, '
        'p99 latency exceeding 2x baseline for more than 10 minutes, or any data integrity '
        'violation detected by continuous validation checks.'
    )
    doc.add_heading('8.1 DNS Failover', level=2)
    doc.add_paragraph(
        'Route 53 weighted routing policies will be configured to enable gradual traffic '
        'shifting. In the event of rollback, DNS weights will be reverted to 100% on-premises '
        'within 60 seconds. TTL values will be reduced to 30 seconds during migration windows.'
    )
    doc.add_heading('8.2 Data Rollback', level=2)
    doc.add_paragraph(
        'Continuous replication from cloud databases back to on-premises will be maintained '
        'during the transition period. Point-in-time recovery capabilities will be enabled '
        'for all Aurora databases with a 7-day retention window.'
    )

    add_page_break(doc)

    # --- Page 11: Security ---
    doc.add_heading('9. Security Considerations', level=1)
    doc.add_paragraph(
        'Security architecture follows a zero-trust model with defense-in-depth. Network '
        'segmentation using VPC security groups and NACLs replaces the current firewall '
        'rules. All data in transit is encrypted using TLS 1.3, and data at rest uses '
        'AWS KMS with customer-managed keys (CMKs).'
    )
    doc.add_heading('9.1 Compliance', level=2)
    doc.add_paragraph(
        'SOC 2 Type II compliance will be maintained throughout the migration. PCI-DSS '
        'scope will be reduced by leveraging AWS-managed services. GDPR data residency '
        'requirements will be enforced through SCPs preventing resource creation outside '
        'eu-central-1 for European customer data.'
    )
    doc.add_heading('9.2 Monitoring and Incident Response', level=2)
    doc.add_paragraph(
        'CloudTrail logs will be centralized in the security account. GuardDuty, Security '
        'Hub, and Macie will provide continuous threat detection. The incident response '
        'playbook will be updated to include cloud-specific scenarios and integrated with '
        'our existing PagerDuty escalation policies.'
    )

    add_page_break(doc)

    # --- Page 12: Cost Analysis ---
    doc.add_heading('10. Cost Analysis and Budget', level=1)

    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    h3 = ['Category', 'Year 1', 'Year 2-3']
    for i, h in enumerate(h3):
        table3.cell(0, i).text = h
        for run in table3.cell(0, i).paragraphs[0].runs:
            run.bold = True
    cost_data = [
        ['AWS Infrastructure', '$540,000', '$480,000/yr'],
        ['Tooling & Licensing', '$85,000', '$45,000/yr'],
        ['Training & Consulting', '$55,000', '$15,000/yr'],
        ['Parallel Operation', '$180,000', 'N/A'],
        ['Contingency (15%)', '$129,000', '$81,000/yr'],
        ['Total', '$989,000', '$621,000/yr'],
    ]
    for r, row_data in enumerate(cost_data, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    doc.add_paragraph(
        'Current on-premises operational cost is $780,000 per year. After full migration, '
        'projected cloud costs are $621,000 per year, representing a 20.4% reduction. '
        'Break-even point is projected at month 22 of the migration timeline.'
    )

    doc.add_heading('11. Timeline and Milestones', level=1)
    milestones = [
        'Q2 2026: Phase 1 complete - Cloud foundation ready',
        'Q3 2026: Phase 2 complete - All databases migrated',
        'Q1 2027: Phase 3 complete - All applications migrated',
        'Q2 2027: Phase 4 complete - Validation and on-premises decommission',
    ]
    for m in milestones:
        doc.add_paragraph(m, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
