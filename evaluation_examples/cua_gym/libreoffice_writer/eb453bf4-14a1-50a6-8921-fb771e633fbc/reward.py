"""
FINAL REWARD SCRIPT - SUCCESS
Task: I just pulled a 40-page draft into LibreOffice Writer 7.6, and every ellipsis is written out as three separate periods ("..."). The copy-editor wants the proper Unicode character instead. How can I run a single Find & Replace that targets the exact string "..." and swaps it for the single ellipsis character "…" across the whole file?
Generated: 2025-09-10 17:35:57
Status: success
Model: azure-o3
Total Steps: 2
"""

import os
import re
from docx import Document

# -------------------------------------------------------------
# Reward script for the LibreOffice Writer ellipsis replacement
# -------------------------------------------------------------
# Task to verify:
#   All occurrences of three consecutive periods "..." should be
#   replaced with a single Unicode ellipsis character "…".
#
# Verification strategy:
#   1. Extract ALL visible text from the DOCX (body, tables,
#      headers, footers).
#   2. Count occurrences of the literal string "..." that are not
#      part of a longer run of periods (regex with look-arounds).
#   3. Count occurrences of the Unicode ellipsis character "…".
#   4. Award points proportionally:             
#        score = ellipsis_count / (ellipsis_count + triple_count)
#      • 1.0 if no "..." remain and at least one "…" exists.
#      • 0.0 if only "..." exist and no ellipsis inserted.
#      • Progressive values in between for partial completion.
#   5. Print detailed diagnostics so the user understands the
#      outcome.
# -------------------------------------------------------------

DOC_PATH = (
    "/home/user/"
    "i_just_pulled_a_40_page_draft_into_libreoffice_writer_76_" 
    "and_every_ellipsis_is_written_out_as_three_.docx"
)


def extract_all_text(doc_path):
    """Return a single string containing all visible text in the doc."""
    doc = Document(doc_path)
    collected = []

    # Body paragraphs
    for p in doc.paragraphs:
        collected.append(p.text)

    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                collected.append(cell.text)

    # Headers & Footers for every section
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                collected.append(p.text)
        if section.footer:
            for p in section.footer.paragraphs:
                collected.append(p.text)

    return "\n".join(collected)


def verify_ellipsis_replacement(file_path: str) -> float:
    """Verify ellipsis replacement and return a score between 0-1."""
    print(f"Checking file: {file_path}")

    # -------- Prerequisite check (no points awarded) --------
    if not os.path.exists(file_path):
        print("✗ File not found – task cannot be verified")
        print("REWARD: 0.0")
        return 0.0

    # -------- Extract text --------
    try:
        all_text = extract_all_text(file_path)
    except Exception as exc:
        print(f"✗ Error reading document: {exc}")
        print("REWARD: 0.0")
        return 0.0

    # -------- Count occurrences --------
    # Regex: exactly three dots not preceded/followed by a dot.
    triple_regex = re.compile(r"(?<!\.)\.\.\.(?!\.)")
    triple_count = len(triple_regex.findall(all_text))
    ellipsis_count = all_text.count("…")

    print(f"Occurrences of literal '...': {triple_count}")
    print(f"Occurrences of Unicode ellipsis '…': {ellipsis_count}")

    # -------- Scoring logic --------
    if triple_count == 0 and ellipsis_count == 0:
        # Edge case: nothing to replace (unlikely in this task)
        score = 0.0
        print("✗ Neither '...' nor '…' found – cannot confirm task completion")
    elif triple_count == 0:
        # Perfect – all triples gone, at least one ellipsis present
        score = 1.0
        print("✓ All triples replaced with ellipsis characters")
    else:
        # Partial completion – some triples remain
        total = triple_count + ellipsis_count
        score = ellipsis_count / total  # progressive score 0–1
        print("⚠ Partial replacement – some '...' remain")

    # Round to two decimals for stability, cap at 1.0
    score = round(min(score, 1.0), 2)

    print(f"REWARD: {score}")
    return score


# ---------------- Run verification -----------------
if __name__ == "__main__":
    verify_ellipsis_replacement(DOC_PATH)

