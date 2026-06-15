"""
FINAL REWARD SCRIPT - SUCCESS
Task: Center the "Introduction" heading in this document.
Generated: 2025-10-14 10:43:27
Status: success
Model: azure-o3
Total Steps: 1
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# -------------------------------------------------------------
# Reward Script: Verify that the "Introduction" heading is centered
# -------------------------------------------------------------
# Scoring Logic (Progressive):
#   +0.5  The document contains a paragraph whose text is exactly "Introduction"
#   +0.5  Every one of those paragraphs is CENTER-aligned (directly or via style)
#   +0.25 Partial credit if some (but not all) of the "Introduction" paragraphs are centered
#   → Maximum possible score: 1.0
# -------------------------------------------------------------

def _effective_alignment(paragraph):
    """Return the effective alignment for a paragraph, considering direct formatting
    and style inheritance. Defaults to LEFT if unspecified (Word default)."""
    # Direct paragraph alignment overrides style alignment if set
    direct_align = paragraph.paragraph_format.alignment
    if direct_align is not None:
        return direct_align

    # Fallback: alignment defined in the paragraph's style (if any)
    try:
        style_align = paragraph.style.paragraph_format.alignment
        if style_align is not None:
            return style_align
    except Exception:
        pass  # Style might not have alignment info

    return WD_ALIGN_PARAGRAPH.LEFT  # Word's implicit default alignment


def verify_center_introduction(file_path: str) -> float:
    """Verify that every paragraph whose text is exactly 'Introduction' is centered."""
    print(f"Verifying 'Introduction' heading centering in: {file_path}")

    score = 0.0

    # -------- Step 1: File existence --------
    if not os.path.exists(file_path):
        print("✗ File not found")
        return 0.0  # No points if file is missing

    # -------- Step 2: Load the DOCX --------
    try:
        doc = Document(file_path)
        print(f"✓ Loaded DOCX with {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Failed to load document: {e}")
        return 0.0  # Cannot continue without loading

    # -------- Step 3: Locate 'Introduction' paragraphs --------
    intro_paragraphs = [p for p in doc.paragraphs if p.text.strip().lower() == "introduction"]
    if not intro_paragraphs:
        print("✗ No paragraph with text 'Introduction' found")
        return 0.0  # Requirement unmet → 0 points

    print(f"✓ Found {len(intro_paragraphs)} paragraph(s) with text 'Introduction'")
    score += 0.5  # Presence of heading earns half the points

    # -------- Step 4: Verify centering --------
    centered_count = 0
    for idx, paragraph in enumerate(intro_paragraphs, start=1):
        align = _effective_alignment(paragraph)
        align_readable = {
            WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
            WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
            WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        }.get(align, str(align))
        print(f"  Paragraph {idx} alignment: {align_readable}")
        if align == WD_ALIGN_PARAGRAPH.CENTER:
            centered_count += 1

    if centered_count == len(intro_paragraphs):
        print("✓ All 'Introduction' paragraphs are CENTER-aligned")
        score += 0.5  # Remaining full credit
    elif centered_count > 0:
        partial = 0.25
        print(f"✓ {centered_count}/{len(intro_paragraphs)} paragraph(s) are centered (+{partial} points)")
        score += partial  # Partial credit
    else:
        print("✗ None of the 'Introduction' paragraphs are centered")

    final_score = min(score, 1.0)
    print(f"REWARD: {final_score}")
    return final_score


if __name__ == "__main__":
    DOC_PATH = "/home/user/center_the_introduction_heading_in_this_document.docx"
    verify_center_introduction(DOC_PATH)

