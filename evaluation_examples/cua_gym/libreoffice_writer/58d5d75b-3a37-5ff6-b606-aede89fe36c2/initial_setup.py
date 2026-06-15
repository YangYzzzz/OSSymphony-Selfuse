"""
Initial Setup: Appellate brief with 6 cited cases, no table of authorities
Task ID: writer_legal_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_para(doc, text, level=1, bold=True, size=14, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
    """Add a styled heading paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = alignment
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return para


def add_body_para(doc, text, first_indent=True, space_after=6):
    """Add a body paragraph with Times New Roman 12pt."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return para


def add_case_citation(para, case_name, rest_text):
    """Add a case citation with the case name in italics."""
    run_case = para.add_run(case_name)
    run_case.italic = True
    run_case.font.name = "Times New Roman"
    run_case.font.size = Pt(12)
    run_rest = para.add_run(rest_text)
    run_rest.font.name = "Times New Roman"
    run_rest.font.size = Pt(12)


def add_page_break(doc):
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    from docx.oxml.ns import qn
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # Page setup - standard legal brief format
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ===== PAGE 1: COVER / CAPTION =====
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    add_heading_para(doc, "IN THE COURT OF APPEALS", level=1, size=14)
    add_heading_para(doc, "FOR THE THIRTEENTH CIRCUIT", level=1, size=14)
    add_heading_para(doc, "No. 2025-CA-04892", level=1, size=12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("RIVERSIDE COMMUNITY COALITION,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Appellant,")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_heading_para(doc, "v.", size=12)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("PACIFIC NORTHWEST ENERGY CORPORATION,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Appellee.")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("BRIEF OF APPELLANT")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("Margaret A. Thornton, Esq.\nBar No. 187452\nThornton & Associates LLP\n1200 Commerce Boulevard, Suite 400\nPortland, OR 97204\n(503) 555-0147\nmthornton@thorntonlaw.com\n\nCounsel for Appellant")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ===== PAGE 2: TABLE OF CONTENTS =====
    add_page_break(doc)
    add_heading_para(doc, "TABLE OF CONTENTS", size=14)

    toc_items = [
        ("TABLE OF AUTHORITIES", "ii"),
        ("STATEMENT OF JURISDICTION", "1"),
        ("STATEMENT OF THE ISSUES", "2"),
        ("STATEMENT OF THE CASE", "3"),
        ("SUMMARY OF THE ARGUMENT", "5"),
        ("ARGUMENT", "6"),
        ("I. THE TRIAL COURT ERRED IN GRANTING SUMMARY JUDGMENT", "6"),
        ("II. THE ENVIRONMENTAL IMPACT ASSESSMENT WAS INADEQUATE", "9"),
        ("III. THE COMMUNITY'S DUE PROCESS RIGHTS WERE VIOLATED", "11"),
        ("CONCLUSION", "14"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(f"{item}\t{page}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # ===== PAGE 3: STATEMENT OF JURISDICTION =====
    add_page_break(doc)
    add_heading_para(doc, "STATEMENT OF JURISDICTION", size=14)

    add_body_para(doc, "This Court has jurisdiction pursuant to 28 U.S.C. § 1291. The District Court entered final judgment on November 15, 2024. Appellant filed a timely notice of appeal on December 12, 2024, within the thirty-day period prescribed by Federal Rule of Appellate Procedure 4(a)(1)(A).")

    add_body_para(doc, "The District Court had subject matter jurisdiction under 28 U.S.C. § 1331, as the claims arise under federal environmental law, specifically the National Environmental Policy Act (\"NEPA\"), 42 U.S.C. §§ 4321-4347, and the Due Process Clause of the Fourteenth Amendment.")

    # ===== STATEMENT OF THE ISSUES =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    add_heading_para(doc, "STATEMENT OF THE ISSUES", size=14)

    add_body_para(doc, "1. Whether the trial court erred in granting summary judgment to Appellee when genuine disputes of material fact existed regarding the environmental impact of the proposed energy facility.")
    add_body_para(doc, "2. Whether the environmental impact assessment prepared by Appellee's consultants satisfied the requirements of NEPA when it failed to consider cumulative impacts on local water resources.")
    add_body_para(doc, "3. Whether the community's procedural due process rights were violated when the public comment period was reduced from sixty to fifteen days without adequate notice.")

    # ===== PAGE 3-4: STATEMENT OF THE CASE (cites Brown v. Board p.4) =====
    add_page_break(doc)
    add_heading_para(doc, "STATEMENT OF THE CASE", size=14)

    # Page 3 content - cite Smith v. Jones here
    p = add_body_para(doc, "Riverside Community Coalition (\"the Coalition\") is a nonprofit organization representing over 2,300 residents of Riverside County who will be directly affected by the construction and operation of a 450-megawatt natural gas power plant proposed by Pacific Northwest Energy Corporation (\"PNE\"). The proposed facility would be located within 1.2 miles of three residential neighborhoods and two elementary schools.")

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("As established in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Smith v. Jones", ", 547 U.S. 398, 412 (2006)")
    run = p.add_run(", organizations with members who face concrete environmental harm have standing to challenge agency actions that threaten those interests. The Coalition's members include homeowners, parents of schoolchildren, and small business owners who will bear the direct consequences of increased air pollution, water contamination, and noise levels.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Page 4 content - cite Brown v. Board
    add_body_para(doc, "On March 8, 2023, PNE submitted its application for a construction permit to the State Environmental Quality Commission (\"SEQC\"). The application included a 340-page environmental impact assessment (\"EIA\") prepared by Greenfield Environmental Consulting, a firm retained and paid by PNE. The Coalition submitted detailed written objections during the abbreviated public comment period, raising concerns about groundwater contamination, air quality degradation, and the cumulative effects of three other industrial facilities already operating within a five-mile radius.")

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("The fundamental right to a clean environment has been recognized by courts across jurisdictions. In ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Brown v. Board", ", 612 F.3d 445, 458 (9th Cir. 2010)")
    run = p.add_run(", the Ninth Circuit held that environmental impact assessments must provide a full and fair discussion of significant environmental impacts, including effects on surrounding communities. The EIA submitted by PNE falls far short of this standard.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ===== PAGE 5: SUMMARY OF ARGUMENT (cite Williams v. State) =====
    add_page_break(doc)
    add_heading_para(doc, "SUMMARY OF THE ARGUMENT", size=14)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("The trial court's grant of summary judgment should be reversed for three independent reasons. First, as the court recognized in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Williams v. State", ", 389 F. Supp. 3d 127, 135 (D. Or. 2019)")
    run = p.add_run(", genuine disputes of material fact preclude summary judgment when reasonable minds could differ on the environmental risks posed by industrial facilities in residential areas. Here, the Coalition presented expert testimony from Dr. Elena Vasquez, a hydrologist with twenty years of experience, directly contradicting PNE's consultants on groundwater contamination risks.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "Second, the environmental impact assessment prepared by PNE's retained consultants was fundamentally inadequate. It failed to analyze cumulative impacts from the three existing industrial facilities within a five-mile radius, omitted any discussion of effects on the local aquifer that supplies drinking water to 15,000 residents, and relied on outdated air quality modeling that did not account for prevailing wind patterns affecting the adjacent school zones.")

    add_body_para(doc, "Third, the SEQC violated the Coalition's due process rights by reducing the public comment period from sixty to fifteen days without providing adequate notice or justification. This truncated timeline prevented the Coalition from obtaining and presenting critical expert analyses that would have revealed the deficiencies in PNE's environmental assessment.")

    # ===== PAGE 6: ARGUMENT I (cite Miller v. County) =====
    add_page_break(doc)
    add_heading_para(doc, "ARGUMENT", size=14)
    add_heading_para(doc, "I. THE TRIAL COURT ERRED IN GRANTING SUMMARY JUDGMENT WHEN GENUINE DISPUTES OF MATERIAL FACT EXISTED", size=12, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("This Court reviews a district court's grant of summary judgment de novo. ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Miller v. County", ", 785 F.3d 1042, 1048 (13th Cir. 2015)")
    run = p.add_run(". Summary judgment is appropriate only when there is no genuine dispute as to any material fact and the movant is entitled to judgment as a matter of law. Fed. R. Civ. P. 56(a). The court must view all evidence in the light most favorable to the nonmoving party and draw all reasonable inferences in that party's favor.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "The Coalition presented substantial evidence creating genuine disputes of material fact on at least four critical issues: (1) the risk of groundwater contamination from the facility's cooling system discharge; (2) the projected increase in particulate matter emissions and their health effects; (3) the adequacy of PNE's proposed noise mitigation measures; and (4) the economic impact on surrounding property values.")

    add_body_para(doc, "Dr. Elena Vasquez, the Coalition's expert hydrologist, submitted a detailed report concluding that PNE's cooling system would discharge thermal effluent at temperatures exceeding state water quality standards, potentially contaminating the Riverside Aquifer. This directly contradicted PNE's expert, Dr. Robert Chang, who opined that the discharge would remain within acceptable limits. Where qualified experts reach opposing conclusions on material scientific questions, summary judgment is inappropriate.")

    # ===== PAGE 7: More Argument I (cite Smith v. Jones again) =====
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Moreover, the standing principles articulated in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Smith v. Jones", ", 547 U.S. at 415")
    run = p.add_run(", confirm that the Coalition's members face injuries that are concrete, particularized, and imminent. Mrs. Patricia Okonkwo, who lives 0.8 miles from the proposed site, testified that she has already experienced decreased property values since PNE announced the project, and her son's pediatrician has expressed concern about the potential health effects of increased particulate matter on children with asthma.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "The trial court dismissed this testimony as speculative, but this characterization was error. Mrs. Okonkwo's testimony was based on a formal property appraisal showing a 12% decline in value and documented medical advice from a licensed physician. This is precisely the type of evidence that creates a genuine dispute of material fact requiring resolution at trial, not on summary judgment.")

    # ===== PAGE 8: Argument (cite Davis v. Corp) =====
    add_page_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("The air quality evidence further underscores the existence of genuine factual disputes. The Coalition's atmospheric scientist, Dr. James Whitfield, modeled particulate matter dispersion using current EPA-approved methods and concluded that PM2.5 concentrations in the Maplewood Elementary School zone would exceed National Ambient Air Quality Standards on at least forty days per year. As the court noted in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Davis v. Corp", ", 891 F.3d 234, 247 (13th Cir. 2018)")
    run = p.add_run(", conflicting expert testimony on the health effects of industrial emissions creates material factual disputes that preclude summary judgment.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "PNE's expert used modeling software that was last updated in 2015 and did not account for the valley terrain that causes temperature inversions trapping pollutants near ground level. The difference between the experts' projected emission levels is not a mere disagreement over methodology; it reflects fundamentally different conclusions about whether the facility will endanger the health of nearby residents, including over 600 schoolchildren.")

    # ===== PAGE 9: ARGUMENT II (cite Williams v. State again) =====
    add_heading_para(doc, "II. THE ENVIRONMENTAL IMPACT ASSESSMENT WAS INADEQUATE UNDER NEPA", size=12, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("NEPA requires federal agencies to take a \"hard look\" at the environmental consequences of proposed actions. 42 U.S.C. § 4332(2)(C). As the court explained in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Williams v. State", ", 389 F. Supp. 3d at 141")
    run = p.add_run(", this obligation is not satisfied by a superficial analysis that ignores known environmental risks or fails to consider the cumulative impact of multiple industrial sources in the same geographic area.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "The EIA prepared by Greenfield Environmental Consulting suffers from three fatal deficiencies. First, it completely ignores the cumulative impact of three existing industrial facilities within a five-mile radius: the Cascade Chemical Processing Plant, the Willamette Paper Mill, and the Columbia Basin Cement Works. Together, these facilities already emit significant quantities of sulfur dioxide, nitrogen oxides, and particulate matter. The EIA analyzes the proposed plant's emissions in isolation, as if it were being built in a pristine environment.")

    # ===== PAGE 10: (cite Garcia v. Fed) =====
    add_page_break(doc)
    add_body_para(doc, "Second, the EIA fails to assess the impact on the Riverside Aquifer, the sole drinking water source for over 15,000 residents. Dr. Vasquez's analysis demonstrates that the proposed facility's wastewater injection wells would be located within the aquifer's recharge zone, creating a direct pathway for contamination. The EIA dismisses this risk in a single paragraph, without any supporting hydrogeological analysis.")

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Third, as recognized in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Garcia v. Fed", ", 723 F.3d 891, 904 (13th Cir. 2013)")
    run = p.add_run(", an EIA must use current scientific methodology and data. The air quality modeling in PNE's EIA relies on the AERMOD dispersion model with 2014 meteorological data, despite the availability of more recent data showing significant changes in local wind patterns due to ongoing climate change. The use of outdated data renders the EIA's air quality projections unreliable.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # ===== PAGE 11: (cite Davis v. Corp again) =====
    add_heading_para(doc, "III. THE COMMUNITY'S DUE PROCESS RIGHTS WERE VIOLATED", size=12, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("The Fourteenth Amendment's Due Process Clause protects the right of affected communities to meaningful participation in government decisions that will substantially affect their property and well-being. The decision in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Davis v. Corp", ", 891 F.3d at 252")
    run = p.add_run(", established that reducing a statutorily mandated comment period without adequate justification violates procedural due process when the affected parties can demonstrate prejudice.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "Here, the SEQC reduced the public comment period from sixty days to fifteen days, citing \"administrative efficiency\" and an unexplained need to \"expedite the permitting process.\" This reduction was announced only through a notice posted on the SEQC's website, without direct notification to the Coalition or any other registered interested party. The abbreviated timeline had concrete, prejudicial effects on the Coalition's ability to prepare and submit meaningful comments.")

    # ===== PAGE 12: (cite Smith v. Jones again) =====
    add_page_break(doc)
    add_body_para(doc, "During the original sixty-day comment period, the Coalition had contracted with Dr. Vasquez to conduct a comprehensive hydrogeological study of the Riverside Aquifer. This study, which required field sampling and laboratory analysis, could not be completed within fifteen days. As a result, the Coalition was forced to submit preliminary comments without the benefit of Dr. Vasquez's full analysis, which ultimately revealed the critical contamination pathway that the EIA had overlooked.")

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("The prejudice to the Coalition is clear and demonstrable. As the Supreme Court recognized in ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Smith v. Jones", ", 547 U.S. at 420")
    run = p.add_run(", procedural protections are meaningless unless they afford affected parties a genuine opportunity to be heard. The fifteen-day comment period was constitutionally inadequate given the complexity of the environmental issues at stake and the technical expertise required to evaluate PNE's EIA.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "The SEQC has offered no legitimate justification for the reduction. \"Administrative efficiency\" is not a constitutionally sufficient basis for curtailing due process rights, particularly when the consequence is the construction of a major industrial facility in a residential area without adequate environmental review. The record shows that the SEQC received only twelve comment submissions during the abbreviated period, compared to an average of forty-seven submissions for comparable projects with standard sixty-day periods.")

    # ===== PAGE 13-14: continued argument and conclusion =====
    add_body_para(doc, "Furthermore, the SEQC's failure to provide direct notification to registered interested parties compounded the due process violation. The Coalition had formally registered with the SEQC as an interested party for all energy facility permit applications in Riverside County. Despite this registration, the SEQC did not send any notification of the shortened comment period to the Coalition or its counsel. The Coalition learned of the abbreviated timeline only through a news article published eight days into the fifteen-day period, leaving it with just seven days to prepare and submit comments on a 340-page EIA.")

    # ===== PAGE 14: CONCLUSION (cite Garcia v. Fed again) =====
    add_page_break(doc)
    add_heading_para(doc, "CONCLUSION", size=14)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("For the foregoing reasons, the Coalition respectfully requests that this Court reverse the trial court's grant of summary judgment and remand this case for trial on the merits. The genuine disputes of material fact identified by the Coalition's experts, the fundamental inadequacies of PNE's environmental impact assessment, and the violation of the community's due process rights each independently warrant reversal. As ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    add_case_citation(p, "Garcia v. Fed", ", 723 F.3d at 910")
    run = p.add_run(", reminds us, \"[t]he purpose of environmental review is not to rubber-stamp predetermined outcomes, but to ensure that decision-makers and the public are fully informed of the consequences of their choices.\"")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_body_para(doc, "The residents of Riverside County deserve no less than a full and fair consideration of the environmental and health risks that the proposed facility would impose on their community. The Coalition urges this Court to vindicate those rights by reversing the judgment below.")

    # Signature block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("Respectfully submitted,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Margaret A. Thornton, Esq.\nBar No. 187452\nThornton & Associates LLP\n1200 Commerce Boulevard, Suite 400\nPortland, OR 97204\n(503) 555-0147\nmthornton@thorntonlaw.com\n\nCounsel for Appellant")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Dated: January 15, 2025")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
