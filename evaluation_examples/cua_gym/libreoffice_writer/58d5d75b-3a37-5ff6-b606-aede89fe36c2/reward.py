"""
Reward Script: Insert a table of authorities for an appellate brief
Task ID: writer_legal_037
Domain: libreoffice_writer
Scoring:
  Component 1: TOA heading exists, centered and bold (0.15)
  Component 2: All 6 cases listed with citations (0.35)
  Component 3: Correct page references for each case (0.20)
  Component 4: Case names formatted in italics (0.15)
  Component 5: Right-aligned dot-leader tab stops (0.15)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_037'

# The 6 expected cases with their citations and page references
EXPECTED_CASES = {
    "Smith v. Jones": {
        "citation_fragment": "547 U.S. 398",
        "pages": ["3", "7", "12"],
    },
    "Williams v. State": {
        "citation_fragment": "389 F. Supp. 3d 127",
        "pages": ["5", "9"],
    },
    "Brown v. Board": {
        "citation_fragment": "612 F.3d 445",
        "pages": ["4"],
    },
    "Davis v. Corp": {
        "citation_fragment": "891 F.3d 234",
        "pages": ["8", "11"],
    },
    "Miller v. County": {
        "citation_fragment": "785 F.3d 1042",
        "pages": ["6"],
    },
    "Garcia v. Fed": {
        "citation_fragment": "723 F.3d 891",
        "pages": ["10", "14"],
    },
}


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def find_toa_section(doc):
    """
    Find the TABLE OF AUTHORITIES section heading paragraph.
    Must be an actual heading (not a TOC entry with a tab + page number).
    Returns (paragraph_index, paragraph) or (None, None).
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Must contain the heading text
        if "TABLE OF AUTHORITIES" in text.upper():
            # Exclude TOC entries which have tab + page number like "TABLE OF AUTHORITIES\tii"
            # The actual heading should be just the title, no tab-separated page ref
            # A TOC entry typically has a tab character followed by a short page ref
            parts = text.split('\t')
            if len(parts) == 1 or (len(parts) == 2 and parts[1].strip() == ''):
                # This is the actual section heading (no tab page ref, or empty after tab)
                return i, para
            # Could also be the heading if the text after tab is not a simple page number
            if len(parts) >= 2:
                after_tab = parts[-1].strip()
                # TOC entries have short page numbers like "ii", "1", "14"
                if len(after_tab) <= 4 and (after_tab.isdigit() or after_tab in ('i', 'ii', 'iii', 'iv', 'v')):
                    continue  # This is a TOC entry, skip
                else:
                    return i, para
    return None, None


def find_case_entries(doc, toa_heading_idx):
    """
    Find paragraphs after the TOA heading that contain case entries.
    Returns list of (paragraph_index, paragraph) for case entry paragraphs.
    Stop when we hit the next major section heading.
    """
    entries = []
    section_headings = [
        "STATEMENT OF JURISDICTION", "STATEMENT OF THE ISSUES",
        "STATEMENT OF THE CASE", "SUMMARY OF THE ARGUMENT",
        "ARGUMENT", "CONCLUSION"
    ]
    for i in range(toa_heading_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        # Stop at next section heading
        if any(text.upper().startswith(h) for h in section_headings):
            break
        entries.append((i, doc.paragraphs[i]))
    return entries


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: TOA heading exists, centered and bold (0.15 points)
    try:
        toa_idx, toa_para = find_toa_section(doc)
        if toa_idx is not None:
            # Check centering
            alignment = toa_para.paragraph_format.alignment
            is_centered = alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
            # Check bold
            is_bold = any(run.font.bold for run in toa_para.runs if run.text.strip())
            if is_centered and is_bold:
                print(f"PASS: Component 1 — TOA heading found at P{toa_idx}, centered and bold (0.15 pts)")
                total_score += 0.15
            elif is_centered or is_bold:
                print(f"PARTIAL: Component 1 — TOA heading found but: centered={is_centered}, bold={is_bold} (0.08 pts)")
                total_score += 0.08
            else:
                print(f"PARTIAL: Component 1 — TOA heading found but not centered or bold (0.05 pts)")
                total_score += 0.05
        else:
            print("FAIL: Component 1 — No TABLE OF AUTHORITIES section heading found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # If no TOA section found, remaining components cannot pass
    if toa_idx is None:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Get case entry paragraphs
    case_entries = find_case_entries(doc, toa_idx)
    entry_texts = [(idx, para.text.strip()) for idx, para in case_entries]

    # Component 2: All 6 cases listed with citations (0.35 points)
    # ~0.058 points per case found with citation
    try:
        cases_found = 0
        cases_matched = {}
        for case_name, info in EXPECTED_CASES.items():
            found = False
            for eidx, (pidx, etext) in enumerate(entry_texts):
                if case_name.lower() in etext.lower() and info["citation_fragment"].lower() in etext.lower():
                    found = True
                    cases_matched[case_name] = eidx
                    break
            if found:
                cases_found += 1
                print(f"  FOUND: {case_name} with citation")
            else:
                print(f"  MISSING: {case_name}")

        if cases_found == 6:
            print(f"PASS: Component 2 — All 6 cases listed with citations (0.35 pts)")
            total_score += 0.35
        elif cases_found > 0:
            partial = round(0.35 * cases_found / 6, 2)
            print(f"PARTIAL: Component 2 — {cases_found}/6 cases found ({partial} pts)")
            total_score += partial
        else:
            print("FAIL: Component 2 — No cases found with citations")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Correct page references for each case (0.20 points)
    try:
        pages_correct = 0
        for case_name, info in EXPECTED_CASES.items():
            expected_pages = info["pages"]
            # Find the entry text for this case
            found_text = None
            for pidx, etext in entry_texts:
                if case_name.lower() in etext.lower():
                    found_text = etext
                    break
            if found_text is None:
                print(f"  PAGE MISS: {case_name} — not found in entries")
                continue

            # Extract page numbers from the text (typically after a tab or at end)
            # Pages appear after the citation, often tab-separated
            parts = found_text.split('\t')
            page_part = parts[-1].strip() if len(parts) > 1 else ""
            # Also try: pages might be after the citation in the same text
            if not page_part:
                # Try to find page numbers at the end of the text
                match = re.search(r'[\d,\s]+$', found_text)
                if match:
                    page_part = match.group().strip()

            # Check if all expected pages are present
            all_found = all(p in page_part for p in expected_pages)
            if all_found:
                pages_correct += 1
                print(f"  PAGES OK: {case_name} — {page_part}")
            else:
                print(f"  PAGES WRONG: {case_name} — expected {expected_pages}, got '{page_part}'")

        if pages_correct == 6:
            print(f"PASS: Component 3 — All page references correct (0.20 pts)")
            total_score += 0.20
        elif pages_correct > 0:
            partial = round(0.20 * pages_correct / 6, 2)
            print(f"PARTIAL: Component 3 — {pages_correct}/6 page refs correct ({partial} pts)")
            total_score += partial
        else:
            print("FAIL: Component 3 — No page references correct")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Case names in italics (0.15 points)
    try:
        italic_count = 0
        for case_name, info in EXPECTED_CASES.items():
            # Find the paragraph for this case
            for pidx, para in case_entries:
                if case_name.lower() in para.text.lower():
                    # Check if the case name part is in an italic run
                    has_italic_name = False
                    for run in para.runs:
                        if run.font.italic and case_name.lower() in run.text.lower():
                            has_italic_name = True
                            break
                        # Case name might be split: e.g., "Smith v. Jones" could be one run
                        # or the case name could be a partial match in an italic run
                        if run.font.italic and run.text.strip():
                            # Check if this italic run contains part of the case name
                            # that looks like a case name (contains "v.")
                            if "v." in run.text or "v " in run.text:
                                has_italic_name = True
                                break
                    if has_italic_name:
                        italic_count += 1
                        print(f"  ITALIC OK: {case_name}")
                    else:
                        print(f"  ITALIC MISS: {case_name} — case name not italic")
                    break

        if italic_count == 6:
            print(f"PASS: Component 4 — All 6 case names in italics (0.15 pts)")
            total_score += 0.15
        elif italic_count > 0:
            partial = round(0.15 * italic_count / 6, 2)
            print(f"PARTIAL: Component 4 — {italic_count}/6 case names italic ({partial} pts)")
            total_score += partial
        else:
            print("FAIL: Component 4 — No case names in italics")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Right-aligned page column with dot leaders (0.15 points)
    try:
        dot_leader_count = 0
        for case_name in EXPECTED_CASES:
            for pidx, para in case_entries:
                if case_name.lower() in para.text.lower():
                    has_right_dots = False
                    for ts in para.paragraph_format.tab_stops:
                        if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                            continue
                        if ts.alignment == WD_TAB_ALIGNMENT.RIGHT and ts.leader == WD_TAB_LEADER.DOTS:
                            has_right_dots = True
                            break
                    if has_right_dots:
                        dot_leader_count += 1
                        print(f"  DOT LEADER OK: {case_name}")
                    else:
                        print(f"  DOT LEADER MISS: {case_name}")
                    break

        if dot_leader_count == 6:
            print(f"PASS: Component 5 — All entries have right-aligned dot leaders (0.15 pts)")
            total_score += 0.15
        elif dot_leader_count > 0:
            partial = round(0.15 * dot_leader_count / 6, 2)
            print(f"PARTIAL: Component 5 — {dot_leader_count}/6 entries with dot leaders ({partial} pts)")
            total_score += partial
        else:
            print("FAIL: Component 5 — No right-aligned dot leader tab stops found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
