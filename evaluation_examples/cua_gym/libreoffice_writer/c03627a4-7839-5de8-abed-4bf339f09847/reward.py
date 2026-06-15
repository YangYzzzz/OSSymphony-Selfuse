"""
Reward Script: Suppress first-page headers in thesis chapters
Task ID: writer_acad_051
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): All sections have different_first_page_header_footer enabled (titlePg)
  Component 2 (0.3): All chapter sections (1-3) have empty first page headers
  Component 3 (0.3): All chapter sections (1-3) still have chapter titles in default headers
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_051'

# Expected chapter titles in default headers for sections 1-3
EXPECTED_CHAPTER_HEADERS = {
    1: "Chapter 1: Introduction to Protein Folding",
    2: "Chapter 2: Computational Methods and Force Fields",
    3: "Chapter 3: Results and Discussion",
}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = list(doc.sections)
    num_sections = len(sections)
    print(f"INFO: Document has {num_sections} sections")

    if num_sections < 4:
        print(f"FAIL: Expected at least 4 sections, found {num_sections}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All sections have different_first_page_header_footer enabled (0.4 points)
    # This is the key task change: titlePg element must be present in each section.
    # In initial_env this is False for all sections, in golden_env it should be True.
    try:
        titlePg_count = 0
        checked_sections = min(num_sections, 4)  # sections 0-3
        for i in range(checked_sections):
            sec = sections[i]
            has_titlePg = sec.different_first_page_header_footer
            print(f"  Section {i}: different_first_page_header_footer = {has_titlePg}")
            if has_titlePg:
                titlePg_count += 1

        if titlePg_count == checked_sections:
            print(f"PASS: Component 1 — All {checked_sections} sections have titlePg enabled (0.4 pts)")
            total_score += 0.4
        elif titlePg_count > 0:
            partial = 0.4 * (titlePg_count / checked_sections)  # proportional partial credit
            print(f"PARTIAL: Component 1 — {titlePg_count}/{checked_sections} sections have titlePg ({partial:.2f} pts)")
            if partial > 0:
                total_score += partial
        else:
            print(f"FAIL: Component 1 — No sections have titlePg enabled")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Chapter sections (1-3) have empty first page headers (0.3 points)
    # The first page header should exist but be empty (suppressed).
    # In initial_env, there is no first page header reference at all.
    # In golden_env, first page header exists with empty text.
    try:
        empty_first_hdr_count = 0
        chapter_sections = [1, 2, 3]
        for idx in chapter_sections:
            if idx >= num_sections:
                continue
            sec = sections[idx]
            # Check if titlePg is enabled first (prerequisite for first page header to matter)
            if not sec.different_first_page_header_footer:
                print(f"  Section {idx}: titlePg not enabled, first page header not active")
                continue

            # Check first page header content
            first_hdr = sec.first_page_header
            first_hdr_text = "".join(p.text for p in first_hdr.paragraphs).strip()
            if first_hdr_text == "":
                print(f"  Section {idx}: First page header is empty (suppressed)")
                empty_first_hdr_count += 1
            else:
                print(f"  Section {idx}: First page header has text: {first_hdr_text!r}")

        if empty_first_hdr_count == len(chapter_sections):
            print(f"PASS: Component 2 — All chapter first page headers are empty/suppressed (0.3 pts)")
            total_score += 0.3
        elif empty_first_hdr_count > 0:
            partial = 0.3 * (empty_first_hdr_count / len(chapter_sections))
            print(f"PARTIAL: Component 2 — {empty_first_hdr_count}/{len(chapter_sections)} chapter first page headers empty ({partial:.2f} pts)")
            if partial > 0:
                total_score += partial
        else:
            print(f"FAIL: Component 2 — No chapter first page headers are empty")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Chapter sections (1-3) still have chapter titles in default headers (0.3 points)
    # The default (non-first-page) header must still contain the chapter title.
    # This verifies the task didn't just remove all headers but properly differentiated them.
    # Note: This check is anchored to Component 1 — it only passes when titlePg is enabled
    # AND the default header has the right content, so it won't award points on initial_env
    # where titlePg is False.
    try:
        correct_default_hdr_count = 0
        for idx in chapter_sections:
            if idx >= num_sections:
                continue
            sec = sections[idx]
            # Gate: only count if titlePg is enabled (the task-introduced change)
            if not sec.different_first_page_header_footer:
                print(f"  Section {idx}: titlePg not enabled, skipping default header check")
                continue

            hdr = sec.header
            hdr_text = "".join(p.text for p in hdr.paragraphs).strip()
            expected = EXPECTED_CHAPTER_HEADERS.get(idx, "")
            if expected and expected in hdr_text:
                print(f"  Section {idx}: Default header contains chapter title: {hdr_text!r}")
                correct_default_hdr_count += 1
            else:
                print(f"  Section {idx}: Default header missing chapter title. Found: {hdr_text!r}, expected: {expected!r}")

        if correct_default_hdr_count == len(chapter_sections):
            print(f"PASS: Component 3 — All chapter default headers contain chapter titles (0.3 pts)")
            total_score += 0.3
        elif correct_default_hdr_count > 0:
            partial = 0.3 * (correct_default_hdr_count / len(chapter_sections))
            print(f"PARTIAL: Component 3 — {correct_default_hdr_count}/{len(chapter_sections)} chapter default headers correct ({partial:.2f} pts)")
            if partial > 0:
                total_score += partial
        else:
            print(f"FAIL: Component 3 — No chapter default headers contain chapter titles")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
