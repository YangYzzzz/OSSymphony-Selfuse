"""
Initial Setup: Create a thesis document with headers on every page (including first pages of chapters).
Task ID: writer_acad_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_051'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_chapter_body(doc, chapter_num, chapter_title, paragraphs_text):
    """Add chapter heading and body paragraphs to fill roughly 2 pages."""
    # Chapter heading
    heading = doc.add_heading(f'Chapter {chapter_num}: {chapter_title}', level=1)
    heading.paragraph_format.page_break_before = False  # section break handles page

    # Body paragraphs
    for text in paragraphs_text:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def set_header_text(section, text):
    """Set the default header (non-first-page header) for a section."""
    header = section.header
    header.is_linked_to_previous = False
    # Clear existing content
    for para in header.paragraphs:
        para.clear()
    p = header.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def create_initial():
    doc = Document()

    # Page setup for default section
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- Title Page (Section 0) ---
    title = doc.add_heading('Computational Approaches to Protein Folding Dynamics\nand Structural Prediction', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(120)

    doc.add_paragraph('')
    author = doc.add_paragraph('by\nDr. Elena Vasquez-Morrison')
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(48)
    for run in author.runs:
        run.font.size = Pt(14)

    dept = doc.add_paragraph('Department of Computational Biology\nUniversity of Cambridge\n2025')
    dept.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept.paragraph_format.space_before = Pt(24)

    sec0 = doc.sections[0]
    sec0.top_margin = Inches(1.5)
    sec0.bottom_margin = Inches(1)
    sec0.left_margin = Inches(1.25)
    sec0.right_margin = Inches(1.25)
    # Title page header
    set_header_text(sec0, 'Computational Approaches to Protein Folding Dynamics')

    # ============================================================
    # Chapter 1 - new section
    # ============================================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec1 = doc.sections[1]
    sec1.top_margin = Inches(1)
    sec1.bottom_margin = Inches(1)
    sec1.left_margin = Inches(1.25)
    sec1.right_margin = Inches(1.25)

    chapter1_paragraphs = [
        'The study of protein folding represents one of the most significant challenges in modern computational biology. Since Anfinsen\'s pioneering thermodynamic hypothesis in 1973, researchers have sought to understand the fundamental principles governing the three-dimensional organization of polypeptide chains. The emergence of machine learning approaches, particularly deep neural networks, has fundamentally transformed our capacity to predict protein structures from amino acid sequences alone.',

        'Molecular dynamics simulations have served as the primary computational tool for investigating protein folding pathways. These simulations integrate Newton\'s equations of motion for each atom in the system, typically employing empirical force fields such as AMBER, CHARMM, or OPLS-AA. Recent advances in GPU computing have extended accessible simulation timescales from nanoseconds to milliseconds, enabling the observation of complete folding events for small proteins.',

        'The free energy landscape framework provides a conceptual foundation for understanding protein folding kinetics. In this paradigm, the folding process is described as a diffusive search on a high-dimensional energy surface, where the native state corresponds to the global free energy minimum. Intermediate states, or metastable configurations, populate local minima and can significantly influence folding rates and pathways.',

        'Enhanced sampling methods, including replica exchange molecular dynamics (REMD), metadynamics, and adaptive biasing force calculations, have been developed to overcome the timescale limitations of conventional molecular dynamics. These techniques accelerate the exploration of conformational space by introducing biasing potentials or by exchanging configurations between simulations conducted at different thermodynamic conditions.',

        'The DeepMind AlphaFold system, released in 2020, achieved unprecedented accuracy in the CASP14 protein structure prediction competition. By combining multiple sequence alignment features with attention-based neural network architectures, AlphaFold demonstrated that end-to-end learning could achieve near-experimental accuracy for many protein targets. Subsequent developments, including AlphaFold2 and RoseTTAFold, have further refined these approaches.',

        'Despite these achievements, significant challenges remain. The prediction of protein dynamics, intrinsically disordered regions, and the effects of post-translational modifications continue to require substantial computational resources and methodological innovation. Furthermore, the accurate prediction of protein-protein interactions and the design of novel protein functions represent active areas of research with implications for drug discovery and synthetic biology.',
    ]
    add_chapter_body(doc, 1, 'Introduction to Protein Folding', chapter1_paragraphs)

    # Set header - same text on ALL pages including first page
    set_header_text(sec1, 'Chapter 1: Introduction to Protein Folding')

    # ============================================================
    # Chapter 2 - new section
    # ============================================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec2 = doc.sections[2]
    sec2.top_margin = Inches(1)
    sec2.bottom_margin = Inches(1)
    sec2.left_margin = Inches(1.25)
    sec2.right_margin = Inches(1.25)

    chapter2_paragraphs = [
        'Force field development constitutes a critical component of molecular simulation methodology. The accuracy of any molecular dynamics simulation is fundamentally limited by the quality of the underlying potential energy function. Classical force fields decompose the total potential energy into bonded terms, including bond stretching, angle bending, and torsional rotation, and nonbonded terms comprising van der Waals interactions and electrostatic contributions.',

        'The parametrization of force fields involves fitting potential energy surfaces to quantum mechanical calculations and experimental observables. The AMBER family of force fields has undergone continuous refinement since its introduction in the 1980s. The ff19SB parameter set, released in 2019, incorporates improved backbone torsion parameters derived from extensive quantum mechanical calculations at the MP2/CBS level of theory.',

        'Polarizable force fields represent a significant advancement over fixed-charge models. By incorporating electronic polarization effects through induced dipoles, Drude oscillators, or fluctuating charge models, these force fields can capture the response of electron distributions to changes in the molecular environment. The AMOEBA force field exemplifies this approach, employing atomic multipoles and explicit polarization through mutual induction.',

        'Coarse-grained models offer an alternative approach to extending the accessible length and time scales of molecular simulations. By grouping multiple atoms into single interaction sites, these models reduce the degrees of freedom and enable the simulation of larger systems over longer timescales. The MARTINI force field has been particularly successful in studying membrane systems, protein aggregation, and large-scale conformational changes.',

        'Machine learning potentials, trained on large datasets of quantum mechanical calculations, have emerged as a promising approach for combining the accuracy of quantum methods with the computational efficiency of empirical force fields. Neural network potentials, Gaussian approximation potentials, and moment tensor potentials have demonstrated remarkable accuracy for a variety of molecular systems, including proteins, nucleic acids, and inorganic materials.',

        'The validation of force fields against experimental data remains essential for establishing their reliability. Nuclear magnetic resonance spectroscopy, X-ray crystallography, small-angle X-ray scattering, and single-molecule fluorescence experiments provide complementary structural and dynamic information that serves as benchmarks for computational predictions.',
    ]
    add_chapter_body(doc, 2, 'Computational Methods and Force Fields', chapter2_paragraphs)

    set_header_text(sec2, 'Chapter 2: Computational Methods and Force Fields')

    # ============================================================
    # Chapter 3 - new section
    # ============================================================
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec3 = doc.sections[3]
    sec3.top_margin = Inches(1)
    sec3.bottom_margin = Inches(1)
    sec3.left_margin = Inches(1.25)
    sec3.right_margin = Inches(1.25)

    chapter3_paragraphs = [
        'The application of enhanced sampling techniques to protein folding studies has yielded valuable insights into folding mechanisms and thermodynamics. Replica exchange molecular dynamics, also known as parallel tempering, remains one of the most widely used enhanced sampling methods. In REMD simulations, multiple copies of the system are simulated simultaneously at different temperatures, and periodic exchange attempts between neighboring replicas are accepted or rejected according to the Metropolis criterion.',

        'Metadynamics, introduced by Laio and Parrinello in 2002, constructs a history-dependent biasing potential that progressively fills free energy minima along selected collective variables. Well-tempered metadynamics improves convergence by gradually reducing the height of deposited Gaussian hills, ensuring that the biasing potential converges to the negative of the free energy surface. This method has been successfully applied to study protein folding, ligand binding, and conformational transitions.',

        'Adaptive biasing force (ABF) methods estimate the mean force along collective variables and apply compensating biases to achieve uniform sampling. The extended adaptive biasing force (eABF) variant decouples the collective variable calculation from the dynamics through an extended Lagrangian formulation, improving numerical stability and reducing artifacts associated with constraint forces.',

        'Transition path sampling and transition interface sampling provide rigorous frameworks for studying rare events without requiring the definition of collective variables a priori. These methods generate an ensemble of reactive trajectories connecting stable states, enabling the characterization of transition state ensembles and the identification of reaction coordinates through committor analysis.',

        'The integration of machine learning with enhanced sampling has opened new possibilities for automated collective variable discovery. Variational autoencoders, time-lagged independent component analysis (TICA), and diffusion maps can identify slow degrees of freedom from simulation trajectories, which can then serve as collective variables for biased simulations. This iterative approach has shown promise for studying complex biomolecular processes.',

        'Markov state models provide a complementary approach for analyzing molecular dynamics trajectories and extracting kinetic information. By discretizing conformational space into metastable states and estimating transition probabilities, these models can predict folding rates, identify intermediate states, and guide the design of additional simulations to improve sampling of underexplored regions of the conformational landscape.',
    ]
    add_chapter_body(doc, 3, 'Results and Discussion', chapter3_paragraphs)

    set_header_text(sec3, 'Chapter 3: Results and Discussion')

    # ============================================================
    # IMPORTANT: In the initial state, headers appear on ALL pages
    # including the first page of each chapter. This means we do NOT
    # set titlePg (different first page) — the same header appears everywhere.
    # ============================================================

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
