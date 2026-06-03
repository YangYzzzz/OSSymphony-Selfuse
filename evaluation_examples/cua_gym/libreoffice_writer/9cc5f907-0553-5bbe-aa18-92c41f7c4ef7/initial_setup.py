"""
Initial Setup: Create sidebar_layout.docx with body text but no text boxes/frames
Task ID: writer_obj_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_028'
OUTPUT = f'{WORKDIR}/sidebar_layout.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)   # 3cm
    section.right_margin = Inches(1.18)  # 3cm
    section.top_margin = Inches(0.98)    # 2.5cm
    section.bottom_margin = Inches(0.98) # 2.5cm

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Urban Planning Report: Downtown Revitalization Project")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    title_para.paragraph_format.space_after = Pt(12)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 1 heading
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Executive Summary")
    h1_run.bold = True
    h1_run.font.size = Pt(13)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(6)

    # Body text paragraphs
    doc.add_paragraph(
        "The Downtown Revitalization Project aims to transform the central business district "
        "of Maplewood City into a vibrant, mixed-use urban center. This initiative has been "
        "developed in collaboration with local stakeholders, including residents, business "
        "owners, and community organizations over the past 18 months."
    )

    doc.add_paragraph(
        "Key objectives include increasing pedestrian traffic by 40%, reducing vehicle "
        "congestion on Main Street by 25%, and attracting at least 15 new businesses "
        "to the area within the next three years. The total projected budget for Phase 1 "
        "is $4.7 million, with phased disbursements planned through fiscal year 2027."
    )

    # Section 2 heading
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Project Background")
    h2_run.bold = True
    h2_run.font.size = Pt(13)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Maplewood City's downtown district experienced a significant economic decline "
        "between 2015 and 2020, with retail vacancy rates rising to 34% and foot traffic "
        "dropping by nearly half compared to 2010 levels. Following extensive community "
        "consultations led by Planning Director Maria Holloway and Deputy Mayor James Griffiths, "
        "the City Council approved the Downtown Revitalization Initiative in October 2023."
    )

    doc.add_paragraph(
        "The project draws on successful models from comparable mid-sized cities, including "
        "the Riverside Arts District transformation in Hartford and the Commerce Row "
        "redevelopment in Burlington. Both projects demonstrated that combining public space "
        "improvements with business incentive programs yields sustainable long-term benefits."
    )

    # Section 3 heading
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Phase 1 Scope of Work")
    h3_run.bold = True
    h3_run.font.size = Pt(13)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Phase 1 covers the six-block stretch of Main Street between Oak Avenue and Elm "
        "Boulevard. Work will include resurfacing 2.4 km of sidewalks, installing 42 new "
        "street trees with protective grates, adding 18 bicycle parking stations, and "
        "upgrading all pedestrian crossing signals to meet current accessibility standards."
    )

    doc.add_paragraph(
        "Streetscape lighting will be redesigned using energy-efficient LED fixtures, "
        "reducing annual electricity costs by an estimated $38,000. Public seating will "
        "be increased by 60%, with a mix of permanent benches and moveable café-style "
        "furniture to accommodate seasonal programming and events."
    )

    # Section 4 heading
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Community Engagement")
    h4_run.bold = True
    h4_run.font.size = Pt(13)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Three public open houses were held in February, May, and September 2024, each "
        "attracting over 200 residents. An online survey collected 1,847 responses, with "
        "78% of participants expressing strong support for the project. Business owners on "
        "Main Street were consulted individually by the Economic Development Office."
    )

    doc.add_paragraph(
        "Priority concerns raised by the community included maintaining access to existing "
        "parking, minimizing construction disruption during the holiday shopping season, "
        "and ensuring that public space improvements benefit residents of all income levels. "
        "These concerns have been addressed in the revised project timeline presented herein."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
