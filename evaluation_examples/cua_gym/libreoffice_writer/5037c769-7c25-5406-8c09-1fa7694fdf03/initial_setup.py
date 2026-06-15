"""
Initial Setup: Configure LibreOffice Writer with default track changes settings
Task ID: writer_rm_042
Domain: libreoffice_writer

Creates a simple .docx file and opens it in LibreOffice Writer.
Ensures the track changes display settings are at their defaults:
  - Insertions: Underline (attribute=3), Color: By Author (-1)
  - Deletions: Underline (attribute=3), Color: By Author (-1)
  - Changed Attributes: Bold (attribute=1), Color: By Author (-1)
"""

import os
import re
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_042'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
XCU_PATH = os.path.expanduser('~/.config/libreoffice/4/user/registrymodifications.xcu')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def kill_libreoffice():
    """Kill any running LibreOffice instances."""
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(2)


def reset_track_changes_defaults():
    """Reset track changes display settings to LibreOffice defaults using string manipulation."""
    if not os.path.exists(XCU_PATH):
        print(f'registrymodifications.xcu not found at {XCU_PATH}, defaults assumed')
        return

    with open(XCU_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    # Remove all existing Revision/TextDisplay entries
    # These are <item> elements with paths containing Revision/TextDisplay
    content = re.sub(
        r'<item\s+oor:path="/org\.openoffice\.Office\.Writer/Revision/TextDisplay/[^"]*">'
        r'<prop[^>]*><value>[^<]*</value></prop></item>\s*',
        '',
        content
    )

    # Also remove LinesChanged entries to keep things clean (we'll re-add defaults)
    content = re.sub(
        r'<item\s+oor:path="/org\.openoffice\.Office\.Writer/Revision/LinesChanged">'
        r'<prop[^>]*><value>[^<]*</value></prop></item>\s*',
        '',
        content
    )

    # Now add back the default entries before </oor:items>
    default_entries = """<item oor:path="/org.openoffice.Office.Writer/Revision/TextDisplay/Insert"><prop oor:name="Attribute" oor:op="fuse"><value>3</value></prop></item>
<item oor:path="/org.openoffice.Office.Writer/Revision/TextDisplay/Insert"><prop oor:name="Color" oor:op="fuse"><value>-1</value></prop></item>
<item oor:path="/org.openoffice.Office.Writer/Revision/TextDisplay/Delete"><prop oor:name="Attribute" oor:op="fuse"><value>3</value></prop></item>
<item oor:path="/org.openoffice.Office.Writer/Revision/TextDisplay/Delete"><prop oor:name="Color" oor:op="fuse"><value>-1</value></prop></item>
<item oor:path="/org.openoffice.Office.Writer/Revision/TextDisplay/ChangedAttribute"><prop oor:name="Attribute" oor:op="fuse"><value>1</value></prop></item>
<item oor:path="/org.openoffice.Office.Writer/Revision/TextDisplay/ChangedAttribute"><prop oor:name="Color" oor:op="fuse"><value>-1</value></prop></item>
<item oor:path="/org.openoffice.Office.Writer/Revision/LinesChanged"><prop oor:name="Mark" oor:op="fuse"><value>3</value></prop></item>
<item oor:path="/org.openoffice.Office.Writer/Revision/LinesChanged"><prop oor:name="Color" oor:op="fuse"><value>0</value></prop></item>
"""

    content = content.replace('</oor:items>', default_entries + '</oor:items>')

    with open(XCU_PATH, 'w', encoding='utf-8') as f:
        f.write(content)

    print('Track changes settings reset to defaults')

    # Verify
    with open(XCU_PATH, 'r', encoding='utf-8') as f:
        verify = f.read()
    for path_part in ['TextDisplay/Insert', 'TextDisplay/Delete', 'TextDisplay/ChangedAttribute']:
        count = verify.count(path_part)
        print(f'  Entries for {path_part}: {count}')


def create_document():
    """Create a simple document with realistic content."""
    from docx import Document

    doc = Document()

    # Title
    doc.add_heading('Quarterly Marketing Strategy Report', level=1)

    # Introduction
    doc.add_paragraph(
        'This report outlines the marketing strategy for Q2 2025, including '
        'key initiatives, budget allocations, and performance targets across '
        'all digital and traditional channels.'
    )

    # Section 1
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'The marketing department achieved a 23% increase in lead generation '
        'during Q1, exceeding the target of 18%. Social media engagement grew '
        'by 31% month-over-month, driven primarily by the new video content '
        'strategy launched in February.'
    )
    doc.add_paragraph(
        'However, the paid advertising cost per acquisition (CPA) rose to $47.80, '
        'up from $38.50 in the previous quarter. This increase is attributed to '
        'heightened competition in the search advertising space and seasonal '
        'fluctuations in ad inventory pricing.'
    )

    # Section 2
    doc.add_heading('Channel Performance', level=2)

    doc.add_heading('Digital Advertising', level=3)
    doc.add_paragraph(
        'Google Ads campaigns generated 4,230 qualified leads at an average '
        'CPA of $42.15. The brand awareness campaigns on display networks '
        'reached 2.8 million unique impressions with a click-through rate of 0.34%.'
    )

    doc.add_heading('Social Media', level=3)
    doc.add_paragraph(
        'Instagram and LinkedIn continue to be our strongest performing platforms. '
        'The LinkedIn thought leadership series attracted 15,600 engagements, '
        'while Instagram Stories achieved an average completion rate of 78%.'
    )

    doc.add_heading('Email Marketing', level=3)
    doc.add_paragraph(
        'The segmented email campaigns delivered an open rate of 28.4% and a '
        'click-through rate of 4.7%, both above industry benchmarks. The nurture '
        'sequence for enterprise prospects showed particularly strong performance '
        'with a 12.3% conversion rate to sales-qualified leads.'
    )

    # Section 3
    doc.add_heading('Budget Recommendations', level=2)
    doc.add_paragraph(
        'Based on the Q1 performance data, we recommend reallocating 15% of the '
        'display advertising budget toward video content production and distribution. '
        'The projected ROI for this shift is estimated at 340% over the next two '
        'quarters, based on the performance of our pilot video campaigns.'
    )

    # Section 4
    doc.add_heading('Next Steps', level=2)
    items = [
        'Finalize the Q2 content calendar by April 15th',
        'Launch the A/B testing framework for landing pages',
        'Conduct quarterly brand perception survey',
        'Review and update customer persona documentation',
        'Negotiate renewed contracts with advertising partners',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial document created: {OUTPUT}')


def main():
    # Kill any running LibreOffice to allow config changes
    kill_libreoffice()

    # Install python-docx if needed
    subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

    # Reset track changes to defaults
    reset_track_changes_defaults()

    # Create the document
    create_document()

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
