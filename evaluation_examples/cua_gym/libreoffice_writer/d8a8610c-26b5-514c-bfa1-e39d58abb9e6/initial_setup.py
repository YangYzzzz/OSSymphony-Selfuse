"""
Initial Setup: Appellate brief document without TOC or Table of Authorities
Task ID: writer_legal_093
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_093'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)

    # ---- Cover / Caption ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run("IN THE UNITED STATES COURT OF APPEALS\nFOR THE NINTH CIRCUIT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()  # blank line

    case_para = doc.add_paragraph()
    case_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = case_para.add_run("No. 24-7831")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    parties_para = doc.add_paragraph()
    parties_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = parties_para.add_run(
        "GREENFIELD ENVIRONMENTAL COALITION,\nPlaintiff-Appellant,\n\nv.\n\n"
        "PACIFICOR ENERGY CORPORATION and\n"
        "UNITED STATES ENVIRONMENTAL PROTECTION AGENCY,\nDefendants-Appellees."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    appeal_para = doc.add_paragraph()
    appeal_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = appeal_para.add_run(
        "On Appeal from the United States District Court\n"
        "for the District of Oregon\n"
        "Case No. 3:22-cv-01547-MKR"
    )
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    brief_title = doc.add_paragraph()
    brief_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = brief_title.add_run("BRIEF OF PLAINTIFF-APPELLANT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    counsel_para = doc.add_paragraph()
    counsel_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = counsel_para.add_run(
        "Elena M. Vasquez\nSenior Litigation Counsel\n"
        "HARTWELL & CHEN LLP\n"
        "1200 SW Morrison Street, Suite 2400\n"
        "Portland, Oregon 97205\n"
        "(503) 555-4200\nevasquez@hartwellchen.com\n\n"
        "Counsel for Plaintiff-Appellant\n"
        "Greenfield Environmental Coalition"
    )
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Page break before body of brief
    doc.add_page_break()

    # ---- PRELIMINARY STATEMENT ----
    h1 = doc.add_heading("PRELIMINARY STATEMENT", level=1)
    for run in h1.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "Plaintiff-Appellant Greenfield Environmental Coalition (\"Greenfield\") "
        "respectfully appeals from the final judgment of the United States District Court "
        "for the District of Oregon, entered on March 15, 2024, granting summary judgment "
        "in favor of Defendants-Appellees PacifiCor Energy Corporation (\"PacifiCor\") and "
        "the United States Environmental Protection Agency (\"EPA\"). The district court erred "
        "in holding that PacifiCor's operation of the Cascade Thermal Generation Station "
        "complied with the Clean Air Act, 42 U.S.C. \u00a7\u00a7 7401-7671q, and that the EPA's "
        "decision to grant a modified Title V operating permit was neither arbitrary nor "
        "capricious under the Administrative Procedure Act, 5 U.S.C. \u00a7 706(2)(A)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # ---- JURISDICTIONAL STATEMENT ----
    h2 = doc.add_heading("JURISDICTIONAL STATEMENT", level=1)
    for run in h2.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The district court had jurisdiction under 28 U.S.C. \u00a7 1331 (federal question) "
        "and 42 U.S.C. \u00a7 7604(a) (citizen suit provision of the Clean Air Act). This Court "
        "has jurisdiction under 28 U.S.C. \u00a7 1291. The district court entered final judgment "
        "on March 15, 2024, and Greenfield filed a timely notice of appeal on April 12, 2024, "
        "within the sixty-day period prescribed by Federal Rule of Appellate Procedure 4(a)(1)(B)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # ---- STATEMENT OF ISSUES PRESENTED ----
    h3 = doc.add_heading("STATEMENT OF ISSUES PRESENTED", level=1)
    for run in h3.runs:
        run.font.name = "Times New Roman"

    issues = [
        "Whether the district court erred in concluding that PacifiCor's emissions "
        "of sulfur dioxide (SO\u2082) and particulate matter (PM\u2082.\u2085) from the Cascade "
        "Thermal Generation Station did not violate the emission limitations established "
        "under 42 U.S.C. \u00a7 7411(a)(1) and Oregon Administrative Rule 340-228-0100.",

        "Whether the EPA acted arbitrarily and capriciously in approving PacifiCor's "
        "modified Title V operating permit without adequate consideration of the cumulative "
        "health impacts on the communities of Millbrook and Cedar Falls, in violation of "
        "Executive Order 12898 and EPA's Environmental Justice Policy.",

        "Whether the district court applied the correct standard of review to the EPA's "
        "permit decision under Chevron U.S.A., Inc. v. Natural Resources Defense Council, "
        "Inc., 467 U.S. 837 (1984), and its progeny."
    ]

    for i, issue in enumerate(issues, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {issue}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # ---- STATEMENT OF THE CASE ----
    h4 = doc.add_heading("STATEMENT OF THE CASE", level=1)
    for run in h4.runs:
        run.font.name = "Times New Roman"

    # Sub-heading
    h4a = doc.add_heading("Factual Background", level=2)
    for run in h4a.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "PacifiCor Energy Corporation operates the Cascade Thermal Generation Station "
        "(\"Cascade Station\"), a 1,340-megawatt coal-fired power plant located in Millbrook, "
        "Oregon. Cascade Station has been in continuous operation since 1978 and is one of "
        "the largest single-source emitters of sulfur dioxide and particulate matter in the "
        "Pacific Northwest. (ER 45-46.)"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "In 2019, PacifiCor applied to the Oregon Department of Environmental Quality "
        "(\"ODEQ\") for a modification to its Title V operating permit to accommodate the "
        "installation of a new combustion optimization system. The modification application "
        "proposed revised emission limits that would allow a temporary increase in SO\u2082 "
        "emissions during a three-year transition period. (ER 112-115.) The EPA, exercising "
        "its oversight authority under 42 U.S.C. \u00a7 7661d(b), reviewed and approved the "
        "modified permit on June 30, 2021. (ER 200-203.)"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The communities of Millbrook and Cedar Falls, located within five miles of "
        "Cascade Station, are predominantly low-income communities with significant minority "
        "populations. As documented in the Greenfield Health Impact Assessment, residents "
        "experience rates of asthma hospitalization 2.7 times the state average and rates of "
        "chronic obstructive pulmonary disease 1.9 times the national average. (ER 301-315.) "
        "A peer-reviewed epidemiological study published in the Journal of Environmental "
        "Health Perspectives found a statistically significant correlation (p < 0.001) "
        "between proximity to Cascade Station and adverse respiratory outcomes. See Rivera "
        "et al., \"Air Quality and Respiratory Health in Communities Adjacent to Coal-Fired "
        "Power Plants,\" 45 J. Envtl. Health Persp. 223, 231-35 (2022)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    h4b = doc.add_heading("Procedural History", level=2)
    for run in h4b.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "On September 3, 2022, Greenfield filed suit against PacifiCor and the EPA in the "
        "United States District Court for the District of Oregon, alleging violations of the "
        "Clean Air Act and the Administrative Procedure Act. Greenfield's complaint asserted "
        "three claims: (1) that PacifiCor's emissions exceeded the New Source Performance "
        "Standards under 42 U.S.C. \u00a7 7411; (2) that the EPA's approval of the modified "
        "permit was arbitrary and capricious under 5 U.S.C. \u00a7 706(2)(A); and (3) that the "
        "EPA failed to comply with Executive Order 12898 regarding environmental justice. "
        "(ER 1-25.)"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "Following discovery, both parties moved for summary judgment. On March 15, 2024, "
        "the Honorable Judge Margaret K. Rodriguez granted Defendants' motion on all three "
        "claims. Greenfield Envtl. Coalition v. PacifiCor Energy Corp., No. 3:22-cv-01547-MKR, "
        "2024 WL 1234567 (D. Or. Mar. 15, 2024). This timely appeal followed."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # ---- SUMMARY OF ARGUMENT ----
    h5 = doc.add_heading("SUMMARY OF ARGUMENT", level=1)
    for run in h5.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The district court committed reversible error in three respects. First, it misapplied "
        "the emission limitation standards under the Clean Air Act by accepting PacifiCor's "
        "averaging methodology, which the statute does not authorize. The plain text of "
        "42 U.S.C. \u00a7 7411(a)(1) requires continuous compliance with emission standards, "
        "not compliance measured through rolling 30-day averages that mask peak exceedances. "
        "See Sierra Club v. Costle, 657 F.2d 298, 319-21 (D.C. Cir. 1981)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "Second, the district court applied an impermissibly deferential standard of review "
        "to the EPA's permit decision. While Chevron deference applies to the EPA's reasonable "
        "interpretation of ambiguous statutory provisions, it does not shield the agency from "
        "scrutiny when it fails to consider an important aspect of the problem or offers an "
        "explanation that runs counter to the evidence. Motor Vehicle Mfrs. Ass'n v. State "
        "Farm Mut. Auto. Ins. Co., 463 U.S. 29, 43 (1983). Here, the EPA entirely failed "
        "to consider the cumulative health impacts on environmental justice communities, "
        "a factor the agency's own policy requires it to evaluate."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "Third, the district court erroneously dismissed Greenfield's environmental justice "
        "claim as non-justiciable. Although Executive Order 12898 does not independently "
        "create a private right of action, the APA provides judicial review of agency action "
        "that is \"arbitrary, capricious, an abuse of discretion, or otherwise not in accordance "
        "with law.\" 5 U.S.C. \u00a7 706(2)(A). The EPA's failure to consider environmental justice "
        "impacts is reviewable under this standard. See Cmtys. Against Runway Expansion, Inc. "
        "v. FAA, 355 F.3d 678, 689 (D.C. Cir. 2004)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # ---- ARGUMENT ----
    h6 = doc.add_heading("ARGUMENT", level=1)
    for run in h6.runs:
        run.font.name = "Times New Roman"

    h6a = doc.add_heading(
        "I. PacifiCor's Emissions Violated the Clean Air Act's New Source Performance Standards",
        level=2
    )
    for run in h6a.runs:
        run.font.name = "Times New Roman"

    h6a1 = doc.add_heading("A. The Standard of Review Is De Novo", level=3)
    for run in h6a1.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "Whether PacifiCor's emissions comply with the applicable New Source Performance "
        "Standards is a question of law that this Court reviews de novo. Nat'l Ass'n of Home "
        "Builders v. Defs. of Wildlife, 551 U.S. 644, 672 (2007). The interpretation of "
        "emission limitation standards under 42 U.S.C. \u00a7 7411 presents a pure question of "
        "statutory construction to which the Court owes no deference to the district court's "
        "analysis."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    h6a2 = doc.add_heading(
        "B. The Plain Language of the Clean Air Act Requires Continuous Compliance",
        level=3
    )
    for run in h6a2.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The Clean Air Act defines \"emission limitation\" as \"a requirement established by "
        "the State or the Administrator which limits the quantity, rate, or concentration of "
        "emissions of air pollutants on a continuous basis.\" 42 U.S.C. \u00a7 7602(k) (emphasis "
        "added). The phrase \"on a continuous basis\" is unambiguous: it requires that emission "
        "limits be met at all times, not merely on average over a 30-day period."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The D.C. Circuit addressed this precise issue in Sierra Club v. Costle, 657 F.2d "
        "298, 319-21 (D.C. Cir. 1981), holding that the EPA could not promulgate emission "
        "standards that permitted averaging periods to mask short-term exceedances. As the "
        "court explained, the \"continuous\" requirement serves a vital protective purpose: it "
        "ensures that communities near polluting facilities are not subjected to intermittent "
        "spikes in harmful emissions, even if long-term averages appear to comply. Id. at 320. "
        "See also Natural Res. Def. Council v. EPA, 489 F.3d 1364, 1373 (D.C. Cir. 2007) "
        "(reaffirming Costle's interpretation)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    h6b = doc.add_heading(
        "II. The EPA's Permit Decision Was Arbitrary and Capricious",
        level=2
    )
    for run in h6b.runs:
        run.font.name = "Times New Roman"

    h6b1 = doc.add_heading(
        "A. The EPA Failed to Consider Cumulative Health Impacts",
        level=3
    )
    for run in h6b1.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "Under the \"hard look\" doctrine, an agency must examine the relevant data and "
        "articulate a satisfactory explanation for its action, including a rational connection "
        "between the facts found and the choice made. Motor Vehicle Mfrs. Ass'n, 463 U.S. at "
        "43. The EPA's permit review failed this standard because the agency did not consider "
        "the cumulative health effects of PacifiCor's emissions on the surrounding communities."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The record demonstrates that the EPA had before it extensive evidence of health "
        "impacts, including the Greenfield Health Impact Assessment and the Rivera et al. "
        "epidemiological study. (ER 301-315, 401-420.) Despite this evidence, the EPA's "
        "permit review focused exclusively on whether PacifiCor's emissions met the numerical "
        "limits in the existing permit, without any analysis of whether those limits adequately "
        "protected public health. This failure to consider an important aspect of the problem "
        "renders the decision arbitrary and capricious. See Pub. Citizen v. Fed. Motor Carrier "
        "Safety Admin., 374 F.3d 1209, 1216 (D.C. Cir. 2004)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    h6b2 = doc.add_heading(
        "B. The EPA's Reliance on Outdated Modeling Was Unreasonable",
        level=3
    )
    for run in h6b2.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The EPA relied on air dispersion modeling from 2015 to assess the impact of "
        "PacifiCor's modified permit, despite the availability of updated modeling techniques "
        "and data. The 2015 model did not account for changes in local meteorological patterns, "
        "population growth in Millbrook and Cedar Falls, or the synergistic effects of "
        "PacifiCor's emissions with those from the nearby Willamette Industrial Complex. "
        "(ER 430-445.) An agency's reliance on stale data in the face of significant changes "
        "renders its decision unreasonable. See Am. Radio Relay League v. FCC, 524 F.3d 227, "
        "236-37 (D.C. Cir. 2008)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    h6c = doc.add_heading(
        "III. The District Court Erred in Dismissing the Environmental Justice Claim",
        level=2
    )
    for run in h6c.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The district court dismissed Greenfield's environmental justice claim on the ground "
        "that Executive Order 12898 does not create a private right of action. While that "
        "premise is correct, the conclusion does not follow. Greenfield does not seek to "
        "enforce Executive Order 12898 directly. Rather, Greenfield challenges the EPA's "
        "failure to consider environmental justice impacts as one factor demonstrating that "
        "the permit decision was arbitrary and capricious under the APA. See 5 U.S.C. "
        "\u00a7 706(2)(A)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "The D.C. Circuit has recognized that an agency's failure to follow its own policies "
        "may inform the arbitrariness inquiry. Cmtys. Against Runway Expansion, Inc. v. FAA, "
        "355 F.3d 678, 689 (D.C. Cir. 2004). The EPA's Environmental Justice Policy, "
        "promulgated pursuant to Executive Order 12898, commits the agency to \"identify and "
        "address\" disproportionate environmental and health effects on minority and low-income "
        "populations. The EPA's complete failure to conduct any environmental justice analysis "
        "in approving PacifiCor's modified permit is precisely the kind of agency oversight "
        "that the APA's arbitrary and capricious standard is designed to catch. See also "
        "Calvert Cliffs' Coordinating Comm. v. U.S. Atomic Energy Comm'n, 449 F.2d 1109, "
        "1115 (D.C. Cir. 1971) (agencies must consider all relevant factors)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # ---- CONCLUSION ----
    h7 = doc.add_heading("CONCLUSION", level=1)
    for run in h7.runs:
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(
        "For the foregoing reasons, Greenfield Environmental Coalition respectfully requests "
        "that this Court reverse the judgment of the district court, vacate the EPA's approval "
        "of PacifiCor's modified Title V operating permit, and remand with instructions for "
        "the EPA to conduct a complete review that includes consideration of cumulative health "
        "impacts and environmental justice concerns."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(
        "\nRespectfully submitted,\n\n"
        "Elena M. Vasquez\n"
        "Senior Litigation Counsel\n"
        "HARTWELL & CHEN LLP\n"
        "1200 SW Morrison Street, Suite 2400\n"
        "Portland, Oregon 97205\n"
        "(503) 555-4200\n"
        "evasquez@hartwellchen.com\n\n"
        "Counsel for Plaintiff-Appellant\n"
        "Greenfield Environmental Coalition\n\n"
        "Dated: April 12, 2024"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
