"""
Initial Setup: Mail merge envelope setup with address spreadsheet
Task ID: writer_lec_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import openpyxl
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_043'
DOCX_OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
XLSX_OUTPUT = f'{WORKDIR}/Desktop/contacts.xlsx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_contacts_spreadsheet():
    """Create contacts.xlsx with 25 US address records."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Contacts"

    headers = ['FirstName', 'LastName', 'Street', 'City', 'State', 'ZIP']
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)

    records = [
        ['Sarah', 'Chen', '1245 Oak Valley Dr', 'Austin', 'TX', '78701'],
        ['Marcus', 'Johnson', '879 Pinehurst Blvd', 'Denver', 'CO', '80202'],
        ['Emily', 'Rodriguez', '3320 Magnolia Ave', 'San Diego', 'CA', '92101'],
        ['David', 'Kim', '556 Birch Creek Ln', 'Portland', 'OR', '97201'],
        ['Jessica', 'Patel', '2100 Riverside Dr', 'Nashville', 'TN', '37201'],
        ['Michael', 'O\'Brien', '410 Sunset Ridge Rd', 'Phoenix', 'AZ', '85001'],
        ['Amanda', 'Torres', '7788 Elm Street', 'Chicago', 'IL', '60601'],
        ['Christopher', 'Nguyen', '1502 Maple Ct', 'Seattle', 'WA', '98101'],
        ['Rachel', 'Washington', '6234 Cedar Hollow Way', 'Atlanta', 'GA', '30301'],
        ['Daniel', 'Martinez', '908 Willow Park Dr', 'Miami', 'FL', '33101'],
        ['Lauren', 'Thompson', '4455 Aspen Trail', 'Minneapolis', 'MN', '55401'],
        ['Andrew', 'Lee', '2678 Spruce Valley Rd', 'Boston', 'MA', '02101'],
        ['Stephanie', 'Garcia', '1833 Hickory Bend Ln', 'Dallas', 'TX', '75201'],
        ['Nathan', 'Clark', '5001 Redwood Cir', 'Sacramento', 'CA', '95814'],
        ['Olivia', 'Wright', '342 Chestnut Hill Ave', 'Philadelphia', 'PA', '19101'],
        ['Tyler', 'Anderson', '7120 Sycamore St', 'Columbus', 'OH', '43201'],
        ['Megan', 'Taylor', '1956 Dogwood Ln', 'Charlotte', 'NC', '28201'],
        ['Brian', 'White', '863 Juniper Ridge Dr', 'Las Vegas', 'NV', '89101'],
        ['Katherine', 'Harris', '4290 Cottonwood Ave', 'San Antonio', 'TX', '78201'],
        ['Jason', 'Robinson', '2115 Poplar Creek Rd', 'Indianapolis', 'IN', '46201'],
        ['Samantha', 'Lewis', '5567 Hawthorn Way', 'Jacksonville', 'FL', '32201'],
        ['Kevin', 'Walker', '3801 Linden Blvd', 'San Francisco', 'CA', '94101'],
        ['Nicole', 'Hall', '1424 Beechwood Dr', 'Baltimore', 'MD', '21201'],
        ['Robert', 'Young', '6752 Cypress Point Ct', 'Houston', 'TX', '77001'],
        ['Jennifer', 'King', '928 Laurel Heights Rd', 'Raleigh', 'NC', '27601'],
    ]

    for r, row_data in enumerate(records, 2):
        for c, val in enumerate(row_data, 1):
            ws.cell(row=r, column=c, value=val)

    # Auto-fit column widths approximately
    ws.column_dimensions['A'].width = 14
    ws.column_dimensions['B'].width = 14
    ws.column_dimensions['C'].width = 24
    ws.column_dimensions['D'].width = 16
    ws.column_dimensions['E'].width = 8
    ws.column_dimensions['F'].width = 8

    wb.save(XLSX_OUTPUT)
    print(f'Contacts spreadsheet created: {XLSX_OUTPUT}')


def create_initial_document():
    """Create a blank Writer document (pre-task state)."""
    doc = Document()

    # Add a simple placeholder paragraph so the document is not completely empty
    para = doc.add_paragraph()
    run = para.add_run('Mail Merge Document')
    run.bold = True
    run.font.size = Pt(16)

    para2 = doc.add_paragraph()
    para2.add_run('This document will be used for mail merge envelope printing. '
                  'Please set up the envelope mail merge using the contacts spreadsheet '
                  'located on the Desktop.')
    para2.paragraph_format.space_after = Pt(12)

    doc.save(DOCX_OUTPUT)
    print(f'Initial document created: {DOCX_OUTPUT}')


def main():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    # Create the contacts spreadsheet
    create_contacts_spreadsheet()

    # Create the initial Writer document
    create_initial_document()

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
