"""
Reward Script: Mail merge envelope setup with address database fields
Task ID: writer_lec_043
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Envelope page dimensions (#10 format: 9.5" x 4.125")
  Component 2 (0.10): Landscape orientation
  Component 3 (0.35): All six address merge fields present in document text
  Component 4 (0.25): Address layout across paragraphs (name / street / city-state-zip)
"""

import os
from docx import Document
from docx.shared import Inches, Emu
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_043'

# #10 envelope dimensions
ENVELOPE_WIDTH_INCHES = 9.5
ENVELOPE_HEIGHT_INCHES = 4.125
TOLERANCE_INCHES = 0.15  # allow small rounding tolerance


def verify_task(file_path):
    """
    Verify mail merge envelope task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: Envelope page dimensions — #10 format 9.5" x 4.125" (0.30 points)
    # The initial doc is standard letter (8.5 x 11), so this only passes on golden.
    try:
        page_w_in = Emu(section.page_width).inches
        page_h_in = Emu(section.page_height).inches

        # Check both orientations: width x height or height x width
        w_match = abs(page_w_in - ENVELOPE_WIDTH_INCHES) < TOLERANCE_INCHES
        h_match = abs(page_h_in - ENVELOPE_HEIGHT_INCHES) < TOLERANCE_INCHES

        # Also allow swapped dimensions (portrait envelope)
        w_match_swap = abs(page_h_in - ENVELOPE_WIDTH_INCHES) < TOLERANCE_INCHES
        h_match_swap = abs(page_w_in - ENVELOPE_HEIGHT_INCHES) < TOLERANCE_INCHES

        if (w_match and h_match) or (w_match_swap and h_match_swap):
            print(f"PASS: Component 1 — Envelope #10 dimensions: {page_w_in:.2f}x{page_h_in:.2f} in (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — Expected ~9.5x4.125 in, found {page_w_in:.2f}x{page_h_in:.2f} in")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Landscape orientation (0.10 points)
    # Initial doc is portrait; envelope should be landscape.
    try:
        if section.orientation == WD_ORIENT.LANDSCAPE:
            print(f"PASS: Component 2 — Landscape orientation (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Expected LANDSCAPE, found {section.orientation}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All six address merge fields present in document text (0.35 points)
    # Initial doc has no merge fields. Golden doc has <FirstName>, <LastName>, etc.
    # We check for the field placeholders in any format: <FieldName>, {FieldName},
    # MERGEFIELD FieldName, or just the field name as text.
    try:
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Also check the raw XML for MERGEFIELD instrText
        body_xml = doc.element.xml

        required_fields = ['FirstName', 'LastName', 'Street', 'City', 'State', 'ZIP']
        fields_found = 0

        for field in required_fields:
            # Check in paragraph text (placeholder style) or in XML (MERGEFIELD style)
            in_text = field in all_text or field.lower() in all_text.lower()
            in_xml = field in body_xml

            if in_text or in_xml:
                fields_found += 1
                print(f"  FOUND field: {field}")
            else:
                print(f"  MISSING field: {field}")

        if fields_found == 6:
            print(f"PASS: Component 3 — All 6 address merge fields present (0.35 pts)")
            total_score += 0.35
        elif fields_found >= 4:
            partial = round(0.35 * (fields_found / 6), 2)
            print(f"PARTIAL: Component 3 — {fields_found}/6 fields found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {fields_found}/6 fields found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Address layout — proper multi-line address format (0.25 points)
    # Golden has 3 paragraphs: name line, street line, city/state/zip line.
    # Initial doc has completely different content (title + instruction).
    # We verify the address is structured across separate paragraphs/lines.
    try:
        all_text = "\n".join(p.text for p in doc.paragraphs)
        body_xml = doc.element.xml

        # Count how many of the 3 address layout lines are present
        layout_parts = 0

        for p in doc.paragraphs:
            ptext = p.text
            pxml = p._element.xml

            # Name line: contains both FirstName and LastName
            if ('FirstName' in ptext or 'FirstName' in pxml) and \
               ('LastName' in ptext or 'LastName' in pxml):
                layout_parts += 1
                break  # count once

        for p in doc.paragraphs:
            ptext = p.text
            pxml = p._element.xml
            # Street line: contains Street
            if 'Street' in ptext or 'Street' in pxml:
                layout_parts += 1
                break

        for p in doc.paragraphs:
            ptext = p.text
            pxml = p._element.xml
            # City/State/ZIP line: contains City and State and ZIP
            city_in = 'City' in ptext or 'City' in pxml
            state_in = 'State' in ptext or 'State' in pxml
            zip_in = 'ZIP' in ptext or 'ZIP' in pxml
            if city_in and state_in and zip_in:
                layout_parts += 1
                break

        if layout_parts == 3:
            print(f"PASS: Component 4 — Correct 3-line address layout (0.25 pts)")
            total_score += 0.25
        elif layout_parts >= 1:
            partial = round(0.25 * (layout_parts / 3), 2)
            print(f"PARTIAL: Component 4 — {layout_parts}/3 address lines correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No recognizable address layout found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
