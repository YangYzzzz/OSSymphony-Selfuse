"""
Initial Setup: Insert a horizontal line separator after the third paragraph on page 1.
Task ID: writer_obj_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_009'
OUTPUT = f'{WORKDIR}/report_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title_para = doc.add_heading('Quarterly Business Performance Report', level=1)

    # 5 paragraphs of realistic text on page 1 — NO horizontal line between them
    paragraphs_text = [
        (
            "This report presents an overview of the company's financial and operational "
            "performance for the third quarter of fiscal year 2025. Overall, the organization "
            "demonstrated resilience in a challenging market environment, achieving a revenue "
            "growth of 8.3% compared to the same period last year. The executive team continues "
            "to monitor key performance indicators closely to ensure alignment with annual targets."
        ),
        (
            "The Sales division recorded its strongest quarter since Q2 2023, driven primarily "
            "by expansion into new regional markets and the successful launch of the Premium "
            "Service tier in July. Account managers across all regions exceeded their quotas, "
            "with the Western Region posting a notable 14.7% year-over-year increase. Customer "
            "acquisition costs decreased by 6% following optimization of the digital marketing "
            "channels in late June."
        ),
        (
            "Operations and supply chain management faced headwinds related to component "
            "shortages, particularly in the electronics manufacturing segment. Despite these "
            "challenges, the logistics team maintained an on-time delivery rate of 94.2%, "
            "surpassing the industry benchmark of 91%. Inventory turnover improved by 1.3 "
            "cycles compared to Q3 2024, reflecting the benefits of the new demand forecasting "
            "system deployed in May."
        ),
        (
            "Human Resources reported a net headcount increase of 47 employees over the "
            "quarter, reaching a total workforce of 1,284 full-time equivalents. Employee "
            "satisfaction scores from the mid-year survey stood at 78 out of 100, a 4-point "
            "improvement from the previous survey. Voluntary turnover remained below the "
            "sector average at 9.1%, attributed to the enhanced benefits package introduced "
            "at the start of the fiscal year."
        ),
        (
            "Looking ahead to Q4, the leadership team has identified three strategic priorities: "
            "accelerating the rollout of the new enterprise software platform, finalizing the "
            "merger integration with Meridian Technologies, and launching the company's "
            "sustainability reporting framework. Budget allocation for these initiatives has "
            "been approved by the Board of Directors and detailed project plans are currently "
            "under review by department heads. Management remains confident in achieving the "
            "full-year guidance communicated to shareholders in February."
        ),
    ]

    for text in paragraphs_text:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
