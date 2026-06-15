"""
Reward Script: Color words in vocabulary worksheet by length
Task ID: osworld_writer_colorword_001
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Words with >6 letters are colored blue (RGB #0000FF)
  Component 2 (0.5): Words with <=6 letters are colored green (RGB #008000)
  Total: 1.0

Task: Color words in a 5x8 table docx:
  - Words with more than 6 letters -> blue text
  - Words with 6 or fewer letters -> green text
  - All 40 cells must be colored (no black text remaining)
"""

import os
from docx import Document
from docx.shared import RGBColor

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_001'

BLUE = (0x00, 0x00, 0xFF)   # Expected color for words > 6 letters
GREEN = (0x00, 0x80, 0x00)  # Expected color for words <= 6 letters
TOLERANCE = 50  # Euclidean RGB distance tolerance for color comparison

# Threshold in the task: "more than 6 letters" = blue, "6 or fewer" = green
# task_config says: >6 letters -> blue, <=6 -> green
# context says: "Words with 7+ letters ... colored blue" (consistent: >6 == >=7)
LONG_THRESHOLD = 6  # word length > LONG_THRESHOLD => blue


def color_distance(rgb_val, target_tuple):
    """Euclidean distance between two RGB triples."""
    r1, g1, b1 = int(rgb_val[0]), int(rgb_val[1]), int(rgb_val[2])
    r2, g2, b2 = target_tuple
    return ((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2) ** 0.5


def get_cell_color(cell):
    """
    Return the RGBColor of the first run with text in the cell,
    or None if no color is set.
    """
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text.strip():
                if run.font.color and run.font.color.rgb is not None:
                    return run.font.color.rgb
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Component 1: All words with >6 letters are colored blue  (0.5 pts)
    Component 2: All words with <=6 letters are colored green (0.5 pts)
    """
    total_score = 0.0

    # Load file — if it fails, nothing can be scored
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have exactly 1 table
    if not doc.tables:
        print("CRITICAL: No tables found in document — expected 1 table with 40 words")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Collect all (word, color) pairs from the table
    cell_results = []  # list of (word_text, rgb_or_none)
    for row in table.rows:
        for cell in row.cells:
            word = cell.text.strip()
            if word:
                color = get_cell_color(cell)
                cell_results.append((word, color))

    if not cell_results:
        print("CRITICAL: Table is empty — no words found")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found {len(cell_results)} words in table")

    # --------------------------------------------------------------------------
    # Component 1: Words with >6 letters are colored blue (0.5 points)
    # --------------------------------------------------------------------------
    try:
        long_words = [(w, c) for (w, c) in cell_results if len(w) > LONG_THRESHOLD]
        if not long_words:
            print("FAIL: Component 1 — no words with >6 letters found (expected some)")
        else:
            blue_ok = []
            blue_fail = []
            for word, color in long_words:
                if color is not None and color_distance(color, BLUE) < TOLERANCE:
                    blue_ok.append(word)
                else:
                    color_str = str(color) if color else 'None/black'
                    blue_fail.append(f"{repr(word)}(color={color_str})")

            print(f"INFO: Component 1 — {len(long_words)} words with >6 letters")
            print(f"  Blue-colored: {len(blue_ok)}/{len(long_words)}")
            if blue_fail:
                # Show up to 5 failures
                print(f"  NOT blue: {blue_fail[:5]}")

            if len(blue_ok) == len(long_words):
                print(f"PASS: Component 1 — all {len(long_words)} long words (>6 letters) are blue (0.5 pts)")
                total_score += 0.5
            else:
                frac = len(blue_ok) / len(long_words)
                print(f"FAIL: Component 1 — only {len(blue_ok)}/{len(long_words)} long words are blue "
                      f"(fraction={frac:.2f}; need all to earn 0.5 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --------------------------------------------------------------------------
    # Component 2: Words with <=6 letters are colored green (0.5 points)
    # --------------------------------------------------------------------------
    try:
        short_words = [(w, c) for (w, c) in cell_results if len(w) <= LONG_THRESHOLD]
        if not short_words:
            print("FAIL: Component 2 — no words with <=6 letters found (expected some)")
        else:
            green_ok = []
            green_fail = []
            for word, color in short_words:
                if color is not None and color_distance(color, GREEN) < TOLERANCE:
                    green_ok.append(word)
                else:
                    color_str = str(color) if color else 'None/black'
                    green_fail.append(f"{repr(word)}(color={color_str})")

            print(f"INFO: Component 2 — {len(short_words)} words with <=6 letters")
            print(f"  Green-colored: {len(green_ok)}/{len(short_words)}")
            if green_fail:
                print(f"  NOT green: {green_fail[:5]}")

            if len(green_ok) == len(short_words):
                print(f"PASS: Component 2 — all {len(short_words)} short words (<=6 letters) are green (0.5 pts)")
                total_score += 0.5
            else:
                frac = len(green_ok) / len(short_words)
                print(f"FAIL: Component 2 — only {len(green_ok)}/{len(short_words)} short words are green "
                      f"(fraction={frac:.2f}; need all to earn 0.5 pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
