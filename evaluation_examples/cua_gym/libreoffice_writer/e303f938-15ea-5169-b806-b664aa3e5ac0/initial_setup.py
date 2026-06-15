"""
Initial Setup: Vocabulary worksheet with table of 40 words, all in default black text
Task ID: osworld_writer_colorword_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title heading
    title = doc.add_heading("Vocabulary Worksheet", level=1)

    # 40 words: mix of short (<=6 letters) and long (>=7 letters)
    # Arranged in 5 rows x 8 columns
    words = [
        # Row 1
        'cat',        'elephant',   'run',        'beautiful',  'big',        'dictionary', 'sun',        'wonderful',
        # Row 2
        'tree',       'umbrella',   'dog',        'adventure',  'hat',        'language',   'cup',        'computer',
        # Row 3
        'pen',        'butterfly',  'map',        'important',  'fox',        'education',  'sky',        'different',
        # Row 4
        'bee',        'mountain',   'ant',        'knowledge',  'mud',        'fantastic',  'red',        'celebrate',
        # Row 5
        'lip',        'creative',   'joy',        'necessary',  'fly',        'together',   'oak',        'possible',
    ]

    # Verify 40 words
    assert len(words) == 40, f"Expected 40 words, got {len(words)}"

    # Create table: 5 rows x 8 columns
    table = doc.add_table(rows=5, cols=8)
    table.style = 'Table Grid'

    for row_idx in range(5):
        for col_idx in range(8):
            word_idx = row_idx * 8 + col_idx
            cell = table.cell(row_idx, col_idx)
            # Clear default empty paragraph and add word run with explicit black color
            para = cell.paragraphs[0]
            run = para.add_run(words[word_idx])
            run.font.size = Pt(12)
            # Explicitly set black color so there's no inherited ambiguity
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
