"""
Initial Setup: Create a document with inconsistent spacing (double paragraph marks)
Task ID: writer_frd_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('Quarterly Performance Review - Q1 2025', level=0)

    # Double paragraph mark 1
    doc.add_paragraph('')

    # --- Section 1: Executive Summary ---
    doc.add_heading('Executive Summary', level=1)

    # Double paragraph mark 2
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'The first quarter of 2025 demonstrated strong growth across multiple business '
        'segments. Total revenue reached $4.2 million, representing a 15% increase over '
        'the same period last year. Our customer acquisition cost decreased by 8%, while '
        'customer lifetime value increased by 12%.'
    )

    # Double paragraph mark 3
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Key achievements during this quarter include the successful launch of our '
        'enterprise platform, expansion into the European market, and the onboarding '
        'of 340 new corporate clients. Employee satisfaction scores remained above 85%, '
        'reflecting our continued investment in workplace culture.'
    )

    # Double paragraph mark 4
    doc.add_paragraph('')

    # --- Section 2: Revenue Analysis ---
    doc.add_heading('Revenue Analysis', level=1)

    # Double paragraph mark 5
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Revenue distribution across our three primary channels showed notable shifts '
        'compared to Q4 2024. Direct sales contributed $2.1 million (50%), partner '
        'channels generated $1.26 million (30%), and digital marketing efforts accounted '
        'for $840,000 (20%).'
    )

    # Double paragraph mark 6
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'The SaaS subscription model continued to outperform expectations with a monthly '
        'recurring revenue of $380,000. Annual contract values for enterprise clients '
        'averaged $45,230, up from $38,900 in the previous quarter. Churn rate remained '
        'stable at 2.3%.'
    )

    # Double paragraph mark 7
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Regional performance varied significantly. North America led with $2.52 million '
        'in revenue, followed by Europe at $1.05 million, and Asia-Pacific at $630,000. '
        'The European market showed the fastest growth rate at 28% quarter-over-quarter.'
    )

    # Double paragraph mark 8
    doc.add_paragraph('')

    # --- Section 3: Team Performance ---
    doc.add_heading('Team Performance', level=1)

    # Double paragraph mark 9
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'The engineering team, led by Sarah Chen, delivered all planned sprint milestones '
        'on schedule. Notable accomplishments include the API gateway redesign, which '
        'reduced average response times by 40%, and the implementation of real-time '
        'analytics dashboards for premium tier customers.'
    )

    # Double paragraph mark 10
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Marcus Johnson and the marketing department executed a comprehensive brand '
        'refresh campaign that resulted in a 25% increase in organic website traffic. '
        'Social media engagement metrics improved across all platforms, with LinkedIn '
        'followers growing by 18,000 during the quarter.'
    )

    # Double paragraph mark 11
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'The customer success team maintained a Net Promoter Score of 72, placing us '
        'in the top quartile of our industry segment. Resolution time for support '
        'tickets averaged 4.2 hours, well within our 6-hour SLA commitment.'
    )

    # Double paragraph mark 12
    doc.add_paragraph('')

    # --- Section 4: Product Development ---
    doc.add_heading('Product Development', level=1)

    # Double paragraph mark 13
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Three major product updates were released during Q1. Version 3.2 introduced '
        'advanced workflow automation features requested by 67% of enterprise clients '
        'in our annual survey. Version 3.3 added multi-language support for 12 additional '
        'languages, expanding our addressable market by an estimated 35%.'
    )

    # Double paragraph mark 14
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'The mobile application underwent a complete redesign based on user research '
        'findings from Dr. Priya Sharma and the UX team. Early adoption metrics show '
        'a 45% increase in daily active users on mobile compared to the previous version. '
        'App store ratings improved from 3.8 to 4.5 stars.'
    )

    # Double paragraph mark 15
    doc.add_paragraph('')

    # --- Section 5: Financial Outlook ---
    doc.add_heading('Financial Outlook', level=1)

    # Double paragraph mark 16
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Based on current trajectory and pipeline analysis, Q2 2025 revenue is projected '
        'at $4.8 million. Operating expenses are expected to increase modestly by 5% due '
        'to planned headcount additions in the engineering and sales departments. EBITDA '
        'margin is forecasted to improve from 18% to 21%.'
    )

    # Double paragraph mark 17
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Strategic investments planned for Q2 include the establishment of a regional '
        'office in Berlin, integration partnerships with three major CRM platforms, and '
        'the launch of our certification program for implementation partners. Total '
        'capital expenditure is budgeted at $620,000.'
    )

    # Double paragraph mark 18
    doc.add_paragraph('')

    p = doc.add_paragraph(
        'Risk factors to monitor include potential currency fluctuations affecting '
        'European revenue, ongoing supply chain constraints impacting hardware delivery '
        'timelines, and increased competitive pressure from two new market entrants '
        'identified in our competitive intelligence reports.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
