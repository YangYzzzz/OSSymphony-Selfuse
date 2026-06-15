"""
Initial Setup: Create study notes document with markdown-style bold markers (**text**)
Task ID: writer_frd_027
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Biology 301 - Study Notes', level=0)

    # Introductory paragraph
    doc.add_paragraph(
        'These notes cover the key topics from the semester. '
        'Terms marked with double asterisks are especially important for the final exam.'
    )

    # Section 1: Cell Biology
    doc.add_heading('Chapter 1: Cell Biology', level=1)
    doc.add_paragraph(
        'The cell is the basic structural and functional unit of all living organisms. '
        'The **plasma membrane** serves as a selective barrier that regulates the passage '
        'of materials into and out of the cell. It is composed of a phospholipid bilayer '
        'embedded with various proteins.'
    )
    doc.add_paragraph(
        'Inside the cell, the **mitochondria** are responsible for generating most of '
        'the cell\'s supply of adenosine triphosphate (ATP), which is used as a source '
        'of chemical energy. They are often referred to as the powerhouses of the cell.'
    )

    # Section 2: Genetics
    doc.add_heading('Chapter 2: Genetics and Heredity', level=1)
    doc.add_paragraph(
        'Gregor Mendel\'s experiments with pea plants established the foundational '
        'principles of genetics. The concept of **dominant and recessive alleles** '
        'explains how traits are inherited across generations. A dominant allele masks '
        'the expression of a recessive allele in heterozygous organisms.'
    )
    doc.add_paragraph(
        'DNA replication is a semiconservative process. The enzyme **DNA polymerase** '
        'synthesizes new strands by reading the template strand in the 3\' to 5\' direction. '
        'Errors during replication are corrected by proofreading mechanisms.'
    )

    # Section 3: Ecology
    doc.add_heading('Chapter 3: Ecology', level=1)
    doc.add_paragraph(
        'Ecosystems are complex networks of interactions between organisms and their '
        'environment. The concept of **trophic levels** organizes organisms by their '
        'feeding position in a food chain: producers, primary consumers, secondary '
        'consumers, and decomposers.'
    )
    doc.add_paragraph(
        'Population dynamics are influenced by birth rates, death rates, immigration, '
        'and emigration. The **carrying capacity** of an environment refers to the '
        'maximum population size that the environment can sustain indefinitely given '
        'available resources.'
    )

    # Section 4: Evolution
    doc.add_heading('Chapter 4: Evolution', level=1)
    doc.add_paragraph(
        'Charles Darwin\'s theory of evolution by **natural selection** proposes that '
        'organisms with favorable traits are more likely to survive and reproduce. Over '
        'many generations, this leads to changes in the characteristics of populations.'
    )
    doc.add_paragraph(
        'Speciation can occur through geographic isolation (allopatric speciation) or '
        'within overlapping populations (sympatric speciation). The study of '
        '**homologous structures** across species provides evidence of common ancestry '
        'and evolutionary divergence.'
    )

    # Section 5: Human Physiology
    doc.add_heading('Chapter 5: Human Physiology', level=1)
    doc.add_paragraph(
        'The human nervous system is divided into the central nervous system (CNS) and '
        'the peripheral nervous system (PNS). Signal transmission at synapses relies on '
        '**neurotransmitters** such as acetylcholine, dopamine, and serotonin, which '
        'bind to receptors on the postsynaptic membrane.'
    )
    doc.add_paragraph(
        'The immune system defends the body against pathogens through innate and '
        'adaptive responses. **Antibodies** are Y-shaped proteins produced by B cells '
        'that specifically bind to antigens, marking them for destruction by other '
        'immune cells.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
