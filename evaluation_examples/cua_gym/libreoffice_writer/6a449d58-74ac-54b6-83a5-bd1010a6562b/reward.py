"""
Reward Script: Convert markdown-style bold markers (**text**) into actual bold formatting
Task ID: writer_frd_027
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): All double-asterisk markers removed from text
  Component 2 (0.5): All 10 expected terms are bold-formatted (0.05 per term)
  Component 3 (0.2): Bold terms are properly inline with surrounding non-bold text
"""

import os
import re

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_027'

# The 10 terms that should be converted from **term** to bold term
EXPECTED_BOLD_TERMS = [
    "plasma membrane",
    "mitochondria",
    "dominant and recessive alleles",
    "DNA polymerase",
    "trophic levels",
    "carrying capacity",
    "natural selection",
    "homologous structures",
    "neurotransmitters",
    "Antibodies",
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather full text and bold runs for analysis
    full_text = ' '.join(p.text for p in doc.paragraphs)
    bold_texts = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.bold is True:
                bold_texts.append(run.text.strip())

    # Component 1: All double-asterisk markers removed (0.3 points)
    # In initial_env there are 10 **term** patterns. In golden there should be 0.
    try:
        asterisk_matches = re.findall(r'\*\*(.+?)\*\*', full_text)
        num_remaining = len(asterisk_matches)
        if num_remaining == 0:
            # Additionally verify no stray double-asterisks remain
            double_ast_count = full_text.count('**')
            if double_ast_count == 0:
                print(f"PASS: Component 1 -- All asterisk markers removed (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 -- {double_ast_count} stray '**' sequences remain")
        else:
            print(f"FAIL: Component 1 -- {num_remaining} asterisk markers still remain: {asterisk_matches}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: All 10 expected terms are bold-formatted (0.5 points, 0.05 per term)
    try:
        terms_found_bold = 0
        for term in EXPECTED_BOLD_TERMS:
            term_lower = term.lower()
            found = False
            for bt in bold_texts:
                if term_lower == bt.lower().strip():
                    found = True
                    break
            if found:
                terms_found_bold += 1
                print(f"PASS: Component 2 -- '{term}' is bold (0.05 pts)")
            else:
                print(f"FAIL: Component 2 -- '{term}' is NOT bold")

        comp2_score = terms_found_bold * 0.05
        total_score += comp2_score
        print(f"Component 2 subtotal: {terms_found_bold}/10 terms bold = {comp2_score:.2f} pts")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Bold terms are properly inline with non-bold surrounding text (0.2 points)
    # This checks that bold runs exist within paragraphs that also have non-bold runs,
    # confirming the bold was applied inline (not entire paragraphs made bold).
    # Fails on initial_env because there are no bold runs at all.
    try:
        inline_bold_count = 0
        for para in doc.paragraphs:
            runs = para.runs
            has_bold = False
            has_nonbold = False
            for run in runs:
                if run.font.bold is True and run.text.strip():
                    has_bold = True
                elif run.text.strip():
                    has_nonbold = True
            if has_bold and has_nonbold:
                inline_bold_count += 1

        # We expect 10 paragraphs with inline bold terms
        # (paragraphs 3,4,6,7,9,10,12,13,15,16 in 0-indexed)
        if inline_bold_count >= 8:
            print(f"PASS: Component 3 -- {inline_bold_count} paragraphs have inline bold formatting (0.2 pts)")
            total_score += 0.2
        elif inline_bold_count >= 5:
            partial = round(0.2 * inline_bold_count / 10, 2)
            print(f"PARTIAL: Component 3 -- {inline_bold_count}/10 paragraphs with inline bold ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 -- Only {inline_bold_count} paragraphs have inline bold formatting")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
