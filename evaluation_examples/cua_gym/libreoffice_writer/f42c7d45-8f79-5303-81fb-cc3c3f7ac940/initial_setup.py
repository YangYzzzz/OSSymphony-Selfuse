"""
Initial Setup: Resignation letter - plain/unformatted docx
Task ID: writer_creative_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'resignation_letter'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Use default Normal style (no explicit formatting) — plain 12pt left-aligned
    # Set default font to 12pt for the document style
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Sender information block (plain, left-aligned, no extra spacing)
    lines = [
        'Thomas Chen',
        '1450 Elm Street, Unit 5',
        'Portland, OR 97201',
        'thomas.chen@email.com',
        '',
        'March 4, 2026',
        '',
        'Ms. Rachel Adams',
        'Vice President of Engineering',
        'TechForward Inc.',
        '200 SW Market Street',
        'Portland, OR 97201',
        '',
        'Dear Ms. Adams,',
        '',
    ]

    for line in lines:
        para = doc.add_paragraph(line)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        for run in para.runs:
            run.font.size = Pt(12)

    # Body paragraph 1
    body1 = (
        'I am writing to formally notify you of my resignation from my position as '
        'Software Engineer at TechForward Inc., effective March 20, 2026. '
        'This was not an easy decision, and I have given it considerable thought.'
    )
    p1 = doc.add_paragraph(body1)
    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    for run in p1.runs:
        run.font.size = Pt(12)

    # Empty line between paragraphs
    sep1 = doc.add_paragraph('')
    sep1.paragraph_format.space_before = Pt(0)
    sep1.paragraph_format.space_after = Pt(0)

    # Body paragraph 2 — contains the key sentence (NOT bolded in initial state)
    body2_part1 = 'My last working day will be March 20, 2026. '
    body2_part2 = (
        'During the remainder of my time here, I am committed to completing my current '
        'projects and assisting with the transition of my responsibilities to ensure '
        'a smooth handover for the team.'
    )
    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run2a = p2.add_run(body2_part1)
    run2a.font.size = Pt(12)
    run2a.bold = False
    run2b = p2.add_run(body2_part2)
    run2b.font.size = Pt(12)
    run2b.bold = False

    # Empty line between paragraphs
    sep2 = doc.add_paragraph('')
    sep2.paragraph_format.space_before = Pt(0)
    sep2.paragraph_format.space_after = Pt(0)

    # Body paragraph 3
    body3 = (
        'I am truly grateful for the opportunities I have had at TechForward Inc. '
        'The experience and skills I have gained here will be invaluable as I move '
        'forward in my career. I look forward to staying in touch and wish the company '
        'continued success.'
    )
    p3 = doc.add_paragraph(body3)
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    for run in p3.runs:
        run.font.size = Pt(12)

    # Empty line before closing
    sep3 = doc.add_paragraph('')
    sep3.paragraph_format.space_before = Pt(0)
    sep3.paragraph_format.space_after = Pt(0)

    # Closing
    closing_lines = [
        'Sincerely,',
        '',
        'Thomas Chen',
    ]
    for line in closing_lines:
        para = doc.add_paragraph(line)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        for run in para.runs:
            run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
