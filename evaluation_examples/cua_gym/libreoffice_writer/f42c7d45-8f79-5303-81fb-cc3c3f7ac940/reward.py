"""
Reward Script: Professional resignation letter formatting
Task ID: writer_creative_018
Domain: libreoffice_writer
Scoring:
  - Component 1: Body paragraphs are justified (0.25 pts)
  - Component 2: Sentence 'My last working day will be March 20, 2026.' is bold (0.30 pts)
  - Component 3: Margins set to 1 inch all around (0.20 pts)
  - Component 4: Font is Liberation Serif on runs that have an explicit font set (0.10 pts)
  - Component 5: PDF exported as ~/Desktop/resignation_letter.pdf (0.15 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_018'

DOCX_PATH = os.path.join(WORKDIR, 'resignation_letter.docx')
PDF_PATH = os.path.join(WORKDIR, 'resignation_letter.pdf')

# The exact sentence that must be bolded
BOLD_SENTENCE = 'My last working day will be March 20, 2026.'


def verify_task(docx_path, pdf_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: ensure the docx file exists and is loadable
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {docx_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Body paragraphs are justified (0.25 points)
    # The three body paragraphs (the content paragraphs) should be JUSTIFY aligned.
    # In initial state they are LEFT aligned. The task asks for justified body paragraphs.
    try:
        # Find body paragraphs: non-empty paragraphs that look like body content
        # Based on exploration: para indices 15, 17, 19 are the three body paragraphs
        body_paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            # Look for body paragraphs: long paragraphs starting with 'I am' or 'My last'
            if (text.startswith('I am writing') or
                    text.startswith('My last working day') or
                    text.startswith('I am truly grateful')):
                body_paragraphs.append(para)

        if len(body_paragraphs) >= 3:
            justified_count = sum(
                1 for p in body_paragraphs
                if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            )
            if justified_count == len(body_paragraphs):
                print(f"PASS: Component 1 — All {len(body_paragraphs)} body paragraphs are JUSTIFY aligned (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — Only {justified_count}/{len(body_paragraphs)} body paragraphs are JUSTIFY aligned")
        else:
            print(f"FAIL: Component 1 — Could not identify {len(body_paragraphs)} body paragraphs (expected 3)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Sentence 'My last working day will be March 20, 2026.' is bold (0.30 points)
    # In initial state run[0] of para 17 has bold=False.
    # In golden state run[0] of para 17 has bold=True.
    try:
        # Collect all runs that contain the bold sentence text
        matching_runs = [
            run
            for para in doc.paragraphs
            for run in para.runs
            if run.text.strip() == BOLD_SENTENCE
        ]

        if len(matching_runs) == 0:
            print(f"FAIL: Component 2 — Could not find run with exact text '{BOLD_SENTENCE}'")
        elif matching_runs[0].bold is True:
            print(f"PASS: Component 2 — '{BOLD_SENTENCE}' is bold=True (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — '{BOLD_SENTENCE}' found but bold={matching_runs[0].bold}, expected True")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Margins set to 1 inch all around (0.20 points)
    # In initial state left/right margins are 1.25 inches, top/bottom are 1.0.
    # The task asks for 1 inch all around.
    try:
        s = doc.sections[0]
        margin_tolerance = 0.05  # allow 0.05 inch tolerance
        left_ok = abs(s.left_margin.inches - 1.0) < margin_tolerance
        right_ok = abs(s.right_margin.inches - 1.0) < margin_tolerance
        top_ok = abs(s.top_margin.inches - 1.0) < margin_tolerance
        bottom_ok = abs(s.bottom_margin.inches - 1.0) < margin_tolerance

        margin_desc = (f"L={s.left_margin.inches:.3f}, R={s.right_margin.inches:.3f}, "
                       f"T={s.top_margin.inches:.3f}, B={s.bottom_margin.inches:.3f}")
        if left_ok and right_ok and top_ok and bottom_ok:
            print(f"PASS: Component 3 — All margins are 1 inch ({margin_desc}) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — Margins not all 1 inch: "
                  f"L={s.left_margin.inches:.3f}, R={s.right_margin.inches:.3f}, "
                  f"T={s.top_margin.inches:.3f}, B={s.bottom_margin.inches:.3f}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Font is Liberation Serif throughout (0.10 points)
    # In initial state runs have font.name=None (inherited).
    # In golden state runs have font.name='Liberation Serif' explicitly set.
    try:
        runs_with_font = []
        total_runs = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    total_runs += 1
                    if run.font.name is not None:
                        runs_with_font.append(run.font.name)

        if total_runs > 0 and len(runs_with_font) > 0:
            lib_serif_count = sum(1 for fn in runs_with_font if fn == 'Liberation Serif')
            # Majority of explicitly-fonted runs should be Liberation Serif
            if lib_serif_count >= len(runs_with_font) * 0.8:
                print(f"PASS: Component 4 — Liberation Serif font set on {lib_serif_count}/{len(runs_with_font)} "
                      f"explicitly-fonted runs out of {total_runs} total (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Only {lib_serif_count}/{len(runs_with_font)} runs use Liberation Serif "
                      f"(expected majority)")
        else:
            print(f"FAIL: Component 4 — No runs with explicit font found (total_runs={total_runs})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: PDF exported as ~/Desktop/resignation_letter.pdf (0.15 points)
    # Initial state: no PDF. Golden state: PDF exists with content.
    try:
        if os.path.exists(pdf_path):
            pdf_size = os.path.getsize(pdf_path)
            if pdf_size > 1000:  # non-trivial size (real PDF)
                print(f"PASS: Component 5 — PDF exists at {pdf_path} (size={pdf_size} bytes) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 5 — PDF exists but is too small ({pdf_size} bytes), may be empty/corrupt")
        else:
            print(f"FAIL: Component 5 — PDF not found at {pdf_path}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Precondition check: ensure docx file exists
if not os.path.exists(DOCX_PATH):
    print(f"File not found: {DOCX_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(DOCX_PATH, PDF_PATH)
