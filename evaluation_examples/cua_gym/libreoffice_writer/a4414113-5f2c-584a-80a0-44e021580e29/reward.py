"""
Reward Script: Add text input field for 'Project Name' and dropdown field for 'Document Type'
Task ID: writer_tech_072
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Text input form field exists in the document
  Component 2 (0.15): Text input field is associated with 'Project Name' context
  Component 3 (0.25): Dropdown form field exists in the document
  Component 4 (0.20): Dropdown field has exactly 4 correct options
  Component 5 (0.15): Dropdown field is associated with 'Document Type' context
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_072'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            import time
            time.sleep(0.8)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Parse all legacy form fields (fldChar with ffData)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body
    fld_chars = body.findall('.//w:fldChar', ns)

    text_input_fields = []
    dropdown_fields = []

    for fc in fld_chars:
        ftype = fc.get(f'{{{ns["w"]}}}fldCharType')
        if ftype != 'begin':
            continue
        ff_data = fc.find('w:ffData', ns)
        if ff_data is None:
            continue

        name_el = ff_data.find('w:name', ns)
        field_name = name_el.get(f'{{{ns["w"]}}}val') if name_el is not None else ''

        has_text_input = ff_data.find('w:textInput', ns) is not None
        dd_list = ff_data.find('w:ddList', ns)
        dd_entries = []
        if dd_list is not None:
            for entry in dd_list.findall('w:listEntry', ns):
                dd_entries.append(entry.get(f'{{{ns["w"]}}}val', ''))

        if has_text_input:
            text_input_fields.append({'name': field_name})
        if dd_list is not None:
            dropdown_fields.append({'name': field_name, 'entries': dd_entries})

    # Also check for SDT-based content controls (alternative implementation)
    sdts = body.findall('.//w:sdt', ns)
    sdt_text_inputs = []
    sdt_dropdowns = []

    for sdt in sdts:
        pr = sdt.find('.//w:sdtPr', ns)
        if pr is None:
            continue
        alias_el = pr.find('w:alias', ns)
        tag_el = pr.find('w:tag', ns)
        alias = alias_el.get(f'{{{ns["w"]}}}val') if alias_el is not None else ''
        tag = tag_el.get(f'{{{ns["w"]}}}val') if tag_el is not None else ''

        has_text = pr.find('w:text', ns) is not None
        dd = pr.find('w:dropDownList', ns)
        has_comboBox = pr.find('w:comboBox', ns) is not None

        if has_text:
            sdt_text_inputs.append({'alias': alias, 'tag': tag})
        if dd is not None:
            items = []
            for item in dd.findall('w:listItem', ns):
                items.append(item.get(f'{{{ns["w"]}}}displayText',
                             item.get(f'{{{ns["w"]}}}value', '')))
            sdt_dropdowns.append({'alias': alias, 'tag': tag, 'items': items})
        if has_comboBox:
            cb = pr.find('w:comboBox', ns)
            items = []
            for item in cb.findall('w:listItem', ns):
                items.append(item.get(f'{{{ns["w"]}}}displayText',
                             item.get(f'{{{ns["w"]}}}value', '')))
            sdt_dropdowns.append({'alias': alias, 'tag': tag, 'items': items})

    has_legacy_text = len(text_input_fields) > 0
    has_sdt_text = len(sdt_text_inputs) > 0
    has_legacy_dd = len(dropdown_fields) > 0
    has_sdt_dd = len(sdt_dropdowns) > 0

    print(f"DEBUG: Legacy text inputs: {text_input_fields}")
    print(f"DEBUG: Legacy dropdowns: {dropdown_fields}")
    print(f"DEBUG: SDT text inputs: {sdt_text_inputs}")
    print(f"DEBUG: SDT dropdowns: {sdt_dropdowns}")

    # Component 1: A text input form field exists (0.25 points)
    try:
        if has_legacy_text or has_sdt_text:
            print(f"PASS: Component 1 — Text input field found (legacy={has_legacy_text}, sdt={has_sdt_text}) (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 1 — No text input form field found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Text input field is associated with 'Project Name' (0.15 points)
    try:
        project_name_found = False
        # Check legacy fields
        for f in text_input_fields:
            if 'project' in f['name'].lower() or 'name' in f['name'].lower():
                project_name_found = True
                break
        # Check SDT fields
        if not project_name_found:
            for f in sdt_text_inputs:
                label = (f['alias'] + ' ' + f['tag']).lower()
                if 'project' in label or 'name' in label:
                    project_name_found = True
                    break
        # Also check: is there a text input field near "Project Name" text in the document?
        if not project_name_found and (has_legacy_text or has_sdt_text):
            # Check if paragraph containing "Project Name" is nearby a form field
            for i, para in enumerate(doc.paragraphs):
                if 'project name' in para.text.lower():
                    # If a text input exists and the paragraph mentions project name,
                    # the field is contextually associated
                    project_name_found = True
                    break

        if project_name_found:
            print(f"PASS: Component 2 — Text input associated with 'Project Name' (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 2 — Text input field not associated with 'Project Name'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: A dropdown form field exists (0.25 points)
    try:
        if has_legacy_dd or has_sdt_dd:
            print(f"PASS: Component 3 — Dropdown field found (legacy={has_legacy_dd}, sdt={has_sdt_dd}) (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 3 — No dropdown form field found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Dropdown has the correct 4 options (0.20 points)
    try:
        expected_options = {'user guide', 'admin guide', 'api reference', 'release notes'}
        dd_correct = False

        # Check legacy dropdown entries
        for f in dropdown_fields:
            actual = {e.lower().strip() for e in f['entries'] if e.strip()}
            if expected_options.issubset(actual):
                dd_correct = True
                print(f"  Found matching legacy dropdown entries: {f['entries']}")
                break

        # Check SDT dropdown items
        if not dd_correct:
            for f in sdt_dropdowns:
                actual = {e.lower().strip() for e in f['items'] if e.strip()}
                if expected_options.issubset(actual):
                    dd_correct = True
                    print(f"  Found matching SDT dropdown items: {f['items']}")
                    break

        if dd_correct:
            print(f"PASS: Component 4 — Dropdown has all 4 required options (0.20 pts)")
            total_score += 0.20
        else:
            all_entries = [f['entries'] for f in dropdown_fields] + [f['items'] for f in sdt_dropdowns]
            print(f"FAIL: Component 4 — Dropdown options don't match. Found: {all_entries}, expected: {expected_options}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Dropdown is associated with 'Document Type' (0.15 points)
    try:
        doc_type_found = False
        # Check legacy fields
        for f in dropdown_fields:
            if 'document' in f['name'].lower() or 'type' in f['name'].lower() or 'doctype' in f['name'].lower():
                doc_type_found = True
                break
        # Check SDT fields
        if not doc_type_found:
            for f in sdt_dropdowns:
                label = (f['alias'] + ' ' + f['tag']).lower()
                if 'document' in label or 'type' in label:
                    doc_type_found = True
                    break
        # Context check: if there's a dropdown and nearby text says "Document Type"
        if not doc_type_found and (has_legacy_dd or has_sdt_dd):
            for i, para in enumerate(doc.paragraphs):
                if 'document type' in para.text.lower():
                    doc_type_found = True
                    break

        if doc_type_found:
            print(f"PASS: Component 5 — Dropdown associated with 'Document Type' (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 5 — Dropdown field not associated with 'Document Type'")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
