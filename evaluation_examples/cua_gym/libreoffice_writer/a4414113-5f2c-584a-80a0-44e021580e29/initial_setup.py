"""
Initial Setup: Technical documentation cover page with static placeholders
Task ID: writer_tech_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_072'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ===== COVER PAGE =====

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Technical Documentation")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.name = "Calibri"

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("CloudSync Platform v3.2")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    run.font.name = "Calibri"

    # Horizontal line (border-bottom on a paragraph)
    line_para = doc.add_paragraph()
    line_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pPr = line_para._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:space'): '1',
        qn('w:color'): '1F497D',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Spacer
    doc.add_paragraph()

    # Project Name - STATIC placeholder (no form field)
    pn_para = doc.add_paragraph()
    pn_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pn_para.paragraph_format.left_indent = Inches(1.5)
    pn_para.paragraph_format.space_after = Pt(12)
    label_run = pn_para.add_run("Project Name: ")
    label_run.bold = True
    label_run.font.size = Pt(12)
    label_run.font.name = "Calibri"
    label_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    value_run = pn_para.add_run("[Enter Project Name]")
    value_run.font.size = Pt(12)
    value_run.font.name = "Calibri"
    value_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Document Type - STATIC placeholder (no form field)
    dt_para = doc.add_paragraph()
    dt_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dt_para.paragraph_format.left_indent = Inches(1.5)
    dt_para.paragraph_format.space_after = Pt(12)
    label_run = dt_para.add_run("Document Type: ")
    label_run.bold = True
    label_run.font.size = Pt(12)
    label_run.font.name = "Calibri"
    label_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    value_run = dt_para.add_run("[Select Document Type]")
    value_run.font.size = Pt(12)
    value_run.font.name = "Calibri"
    value_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Spacer
    doc.add_paragraph()

    # Metadata table on cover page
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_data = [
        ("Author", "Elena Rodriguez"),
        ("Version", "3.2.1"),
        ("Date", "March 28, 2026"),
        ("Department", "Platform Engineering"),
    ]
    for i, (key, val) in enumerate(meta_data):
        cell_key = meta_table.cell(i, 0)
        run_k = cell_key.paragraphs[0].add_run(key)
        run_k.bold = True
        run_k.font.size = Pt(10)
        run_k.font.name = "Calibri"
        cell_val = meta_table.cell(i, 1)
        run_v = cell_val.paragraphs[0].add_run(val)
        run_v.font.size = Pt(10)
        run_v.font.name = "Calibri"

    # Set table column widths
    for row in meta_table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(3.5)

    # Page break after cover page
    doc.add_page_break()

    # ===== PAGE 2: TABLE OF CONTENTS / BODY =====
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    toc_items = [
        "1. Introduction",
        "2. System Architecture",
        "3. API Endpoints",
        "4. Authentication & Authorization",
        "5. Deployment Guide",
        "6. Troubleshooting",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.name = "Calibri"

    doc.add_page_break()

    # ===== PAGE 3: INTRODUCTION =====
    intro_heading = doc.add_heading("1. Introduction", level=1)
    intro_heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    intro_text = (
        "The CloudSync Platform provides a comprehensive set of APIs and services "
        "for real-time data synchronization across distributed systems. This document "
        "covers the core architecture, endpoint specifications, and deployment procedures "
        "required to integrate CloudSync into your existing infrastructure."
    )
    p = doc.add_paragraph(intro_text)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    p2 = doc.add_paragraph(
        "CloudSync supports both synchronous and asynchronous data flows, with built-in "
        "conflict resolution strategies including last-write-wins, vector clocks, and "
        "custom merge functions. The platform handles over 2.4 million requests per second "
        "in production environments with 99.97% uptime."
    )
    p2.paragraph_format.space_after = Pt(8)
    for run in p2.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
