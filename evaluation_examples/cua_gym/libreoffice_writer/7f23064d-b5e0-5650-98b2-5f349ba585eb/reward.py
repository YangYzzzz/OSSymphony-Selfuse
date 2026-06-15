"""
FINAL REWARD SCRIPT - SUCCESS
Task: I’ve got a section titled "Revision" (it’s set as Heading 1). Right underneath that heading I want to place today’s date as a non-updating field and shove it all the way to the right margin. What steps should I follow in LibreOffice Writer to get that done?
Generated: 2025-09-10 18:51:26
Status: success
Model: azure-o3
Total Steps: 9
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# Path to the file that should contain the completed work
FILE_PATH = "/home/user/ive_got_a_section_titled_revision_its_set_as_heading_1_right_underneath_that_heading_i_want_to_place.docx"

def is_date_string(text: str) -> bool:
    """Return True if *text* matches common date formats (e.g., September 10, 2025 or 10/09/2025)."""
    text = text.strip()
    # Written month format (e.g., September 10, 2025)
    month_pattern = r"(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}"
    # Numeric format (e.g., 10/09/2025 or 10-09-25)
    numeric_pattern = r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"
    return bool(re.search(month_pattern, text, re.IGNORECASE) or re.search(numeric_pattern, text))

def paragraph_contains_field(paragraph) -> bool:
    """Detect whether a paragraph contains a Word field (fldChar / instrText)."""
    for element in paragraph._p.iter():  # pylint: disable=protected-access
        tag = element.tag
        if tag.endswith("fldChar") or tag.endswith("instrText"):
            return True
    return False

def verify_writer_task(file_path: str) -> float:
    """Verify that the Writer task has been completed correctly and return a progressive score."""

    print(f"Starting verification for: {file_path}\n")

    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as exc:  # pylint: disable=broad-except
        print(f"✗ Unable to open DOCX file: {exc}")
        print("REWARD: 0.0")
        return 0.0

    total_score = 0.0  # Progressive score

    # ------------------------------------------------------------------
    # 1. Locate "Revision" heading (must be Heading 1)
    # ------------------------------------------------------------------
    revision_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if (
            para.text.strip().lower() == "revision"
            and para.style.name.lower().startswith("heading")
            and "1" in para.style.name
        ):
            revision_idx = idx
            break

    if revision_idx is not None:
        print("✓ Found 'Revision' heading formatted as Heading 1")
        total_score += 0.3
    else:
        print("✗ 'Revision' heading with Heading 1 style not found")

    # ------------------------------------------------------------------
    # 2. Find first non-empty paragraph after the heading & confirm it’s a date
    # ------------------------------------------------------------------
    date_para = None
    if revision_idx is not None:
        for para in doc.paragraphs[revision_idx + 1 :]:
            if para.text.strip():  # skip empty lines
                date_para = para
                break

    if date_para and is_date_string(date_para.text):
        print(f"✓ Found date paragraph immediately after heading: '{date_para.text.strip()}'")
        total_score += 0.3
    else:
        if date_para:
            print(
                f"✗ Paragraph after heading is not recognised as a date: '{date_para.text.strip()}'"
            )
        else:
            print("✗ No non-empty paragraph found after 'Revision' heading")
        # If date wasn't correct, we can't evaluate alignment / field reliably
        final = round(min(total_score, 1.0), 2)
        print(f"REWARD: {final}")
        return final

    # ------------------------------------------------------------------
    # 3. Check that the date paragraph is right-aligned
    # ------------------------------------------------------------------
    if date_para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        print("✓ Date paragraph is right-aligned to the margin")
        total_score += 0.2
    else:
        print("✗ Date paragraph is not right-aligned")

    # ------------------------------------------------------------------
    # 4. Ensure the date is *not* a field (should be plain text)
    # ------------------------------------------------------------------
    if not paragraph_contains_field(date_para):
        print("✓ Date paragraph is plain text (non-updating)")
        total_score += 0.2
    else:
        print("✗ Date paragraph contains an updating field")

    # ------------------------------------------------------------------
    # Final score (capped at 1.0)
    # ------------------------------------------------------------------
    final_score = round(min(total_score, 1.0), 2)
    print(f"\nTotal Score: {final_score} (out of 1.0)")
    print(f"REWARD: {final_score}")
    return final_score

# ----------------------------------------------------------------------
# Run verification when script is executed
# ----------------------------------------------------------------------
if __name__ == "__main__":
    verify_writer_task(FILE_PATH)

