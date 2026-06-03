"""
Initial Setup: Create an invoice document with a 5x8 table with default borders.
Task ID: writer_tm_012
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_012'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    Usage: set_cell_border(cell, top={"sz": 4, "color": "000000", "val": "single"}, ...)
    sz is in eighth-points (4 = 0.5pt, 12 = 1.5pt).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    # Remove existing tcBorders
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('INVOICE', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Invoice metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = meta.add_run('Invoice #: INV-2025-0847\n')
    run.font.size = Pt(11)
    run = meta.add_run('Date: March 15, 2025\n')
    run.font.size = Pt(11)
    run = meta.add_run('Due Date: April 14, 2025')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Bill-to info
    bill_to = doc.add_paragraph()
    run = bill_to.add_run('Bill To: ')
    run.bold = True
    run.font.size = Pt(11)
    run = bill_to.add_run('Meridian Technologies Inc.\n1250 Harbor Boulevard, Suite 400\nSan Diego, CA 92101')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Create 5x8 invoice table (1 header + 7 data rows)
    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid'  # default 0.5pt black borders

    # Headers
    headers = ['Item', 'Description', 'Quantity', 'Unit Price', 'Total']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # Invoice line items - realistic data
    data = [
        ['1', 'Cloud Server Hosting (Monthly)', '3', '$450.00', '$1,350.00'],
        ['2', 'SSL Certificate Renewal (Annual)', '2', '$129.99', '$259.98'],
        ['3', 'Database Migration Service', '1', '$2,800.00', '$2,800.00'],
        ['4', 'Network Security Audit', '1', '$1,750.00', '$1,750.00'],
        ['5', 'API Integration Development', '40', '$95.00', '$3,800.00'],
        ['6', 'Technical Support (Hours)', '15', '$125.00', '$1,875.00'],
        ['7', '24/7 Monitoring Service (Monthly)', '1', '$680.00', '$680.00'],
    ]

    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    doc.add_paragraph()  # spacer

    # Total line
    total_para = doc.add_paragraph()
    total_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = total_para.add_run('Subtotal: $12,514.98\n')
    run.font.size = Pt(11)
    run = total_para.add_run('Tax (8.5%): $1,063.77\n')
    run.font.size = Pt(11)
    run = total_para.add_run('Total Due: $13,578.75')
    run.bold = True
    run.font.size = Pt(13)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
