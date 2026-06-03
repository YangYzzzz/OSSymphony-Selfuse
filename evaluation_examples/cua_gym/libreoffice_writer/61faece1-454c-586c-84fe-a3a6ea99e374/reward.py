"""
Reward Script: Set all borders of invoice table to 1.5pt solid blue (#0000FF) lines
Task ID: writer_tm_012
Domain: libreoffice_writer
Scoring:
  Component 1 (0.45): Table-level borders have blue color (#0000FF) on all 6 border types
  Component 2 (0.35): Table-level border size=12 (1.5pt) and style="single" on all borders
  Component 3 (0.20): Cell-level borders on sampled cells also blue and correct size
  Precondition gate: Content must be intact (not a scoring component)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_012'

# Expected cell content for data integrity precondition
EXPECTED_HEADER = ['Item', 'Description', 'Quantity', 'Unit Price', 'Total']

# Border types to check at table level
TABLE_BORDER_TYPES = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
# Border types to check at cell level
CELL_BORDER_TYPES = ['top', 'left', 'bottom', 'right']

EXPECTED_COLOR = '0000FF'
EXPECTED_SZ = '12'  # 12 eighth-points = 1.5pt
EXPECTED_VAL = 'single'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: must have at least one table with correct dimensions
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition gate: content integrity (NOT a scoring component)
    try:
        if len(table.rows) < 8 or len(table.columns) < 5:
            print(f"CRITICAL: Table dimensions wrong: {len(table.rows)}x{len(table.columns)}, expected 8x5")
            print("REWARD: 0.0")
            return 0.0
        header_row = [table.cell(0, c).text.strip() for c in range(5)]
        if header_row != EXPECTED_HEADER:
            print(f"CRITICAL: Header row corrupted: {header_row}")
            print("REWARD: 0.0")
            return 0.0
        print("PRECONDITION: Table structure and content intact -- proceeding with border checks")
    except Exception as e:
        print(f"CRITICAL: Cannot verify table content: {e}")
        print("REWARD: 0.0")
        return 0.0

    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))

    # Component 1: Table-level borders have blue color (#0000FF) on all 6 border types (0.45 points)
    try:
        tblBorders = None
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))

        if tblBorders is None:
            print("FAIL: Component 1 -- No table-level borders (tblBorders) found")
        else:
            color_pass_count = 0
            for btype in TABLE_BORDER_TYPES:
                elem = tblBorders.find(qn(f'w:{btype}'))
                if elem is not None:
                    color = elem.get(qn('w:color'), '').upper()
                    if color == EXPECTED_COLOR:
                        color_pass_count += 1
                    else:
                        print(f"  INFO: tblBorders/{btype} color={color}, expected={EXPECTED_COLOR}")
                else:
                    print(f"  INFO: tblBorders/{btype} element not found")

            if color_pass_count == len(TABLE_BORDER_TYPES):
                print(f"PASS: Component 1 -- All {len(TABLE_BORDER_TYPES)} table-level borders have blue color #0000FF (0.45 pts)")
                total_score += 0.45
            elif color_pass_count > 0:
                partial = 0.45 * (color_pass_count / len(TABLE_BORDER_TYPES))
                print(f"PARTIAL: Component 1 -- {color_pass_count}/{len(TABLE_BORDER_TYPES)} table borders have blue color ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 -- No table-level borders have blue color")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Table-level border size=12 (1.5pt) and style="single" (0.35 points)
    try:
        tblBorders = None
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))

        if tblBorders is None:
            print("FAIL: Component 2 -- No table-level borders found")
        else:
            sz_style_pass = 0
            for btype in TABLE_BORDER_TYPES:
                elem = tblBorders.find(qn(f'w:{btype}'))
                if elem is not None:
                    sz = elem.get(qn('w:sz'), '')
                    val = elem.get(qn('w:val'), '')
                    if sz == EXPECTED_SZ and val == EXPECTED_VAL:
                        sz_style_pass += 1
                    else:
                        print(f"  INFO: tblBorders/{btype} sz={sz} val={val}, expected sz={EXPECTED_SZ} val={EXPECTED_VAL}")
                else:
                    print(f"  INFO: tblBorders/{btype} element not found")

            if sz_style_pass == len(TABLE_BORDER_TYPES):
                print(f"PASS: Component 2 -- All {len(TABLE_BORDER_TYPES)} table borders are 1.5pt solid (0.35 pts)")
                total_score += 0.35
            elif sz_style_pass > 0:
                partial = 0.35 * (sz_style_pass / len(TABLE_BORDER_TYPES))
                print(f"PARTIAL: Component 2 -- {sz_style_pass}/{len(TABLE_BORDER_TYPES)} table borders are 1.5pt solid ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 -- No table-level borders have correct size/style")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Cell-level borders on sampled cells have blue color and correct size (0.2 points)
    # Sample corner and middle cells to verify cell-level border overrides
    try:
        sample_cells = [(0, 0), (0, 4), (3, 2), (7, 0), (7, 4)]
        cells_with_blue_borders = 0

        for ri, ci in sample_cells:
            if ri >= len(table.rows) or ci >= len(table.columns):
                continue
            cell = table.cell(ri, ci)
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                continue
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                continue

            cell_ok = True
            for btype in CELL_BORDER_TYPES:
                elem = tcBorders.find(qn(f'w:{btype}'))
                if elem is not None:
                    color = elem.get(qn('w:color'), '').upper()
                    sz = elem.get(qn('w:sz'), '')
                    val = elem.get(qn('w:val'), '')
                    if color != EXPECTED_COLOR or sz != EXPECTED_SZ or val != EXPECTED_VAL:
                        cell_ok = False
                        break
            if cell_ok and tcBorders is not None and len(tcBorders) > 0:
                cells_with_blue_borders += 1

        if cells_with_blue_borders >= 3:
            print(f"PASS: Component 3 -- {cells_with_blue_borders}/{len(sample_cells)} sampled cells have blue 1.5pt borders (0.2 pts)")
            total_score += 0.2
        elif cells_with_blue_borders > 0:
            partial = 0.2 * (cells_with_blue_borders / len(sample_cells))
            print(f"PARTIAL: Component 3 -- {cells_with_blue_borders}/{len(sample_cells)} sampled cells have cell-level blue borders ({partial:.2f} pts)")
            total_score += partial
        else:
            # Cell-level borders may not be set if table-level borders suffice.
            # Check if table-level borders pass; if so, give full credit for this component.
            if tblPr is not None:
                tblBorders_check = tblPr.find(qn('w:tblBorders'))
                if tblBorders_check is not None:
                    all_tbl_ok = True
                    for btype in TABLE_BORDER_TYPES:
                        elem = tblBorders_check.find(qn(f'w:{btype}'))
                        if elem is None:
                            all_tbl_ok = False
                            break
                        if (elem.get(qn('w:color'), '').upper() != EXPECTED_COLOR or
                                elem.get(qn('w:sz'), '') != EXPECTED_SZ or
                                elem.get(qn('w:val'), '') != EXPECTED_VAL):
                            all_tbl_ok = False
                            break
                    if all_tbl_ok:
                        print(f"PASS: Component 3 -- No cell-level borders but table-level borders cover all borders correctly (0.2 pts)")
                        total_score += 0.2
                    else:
                        print(f"FAIL: Component 3 -- Neither cell-level nor table-level borders fully configured")
                else:
                    print(f"FAIL: Component 3 -- No cell-level or table-level borders found")
            else:
                print(f"FAIL: Component 3 -- No table properties found")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer (save unsaved GUI edits)
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
