"""
Reward Script: Insert a 4-column, 8-row table with bold header row and bottom border
Task ID: writer_biz_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with 4 columns and 8 rows
  Component 2 (0.25): Header row contains correct text
  Component 3 (0.25): Header row text is bold
  Component 4 (0.25): Header row cells have bottom border
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_030'

def persist_app_state(domain: str):
    """Attempt to save any unsaved GUI state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions (0.25 points)
    # Initial has 0 tables; golden should have at least 1 table with 8 rows and 4 columns
    try:
        if len(doc.tables) < 1:
            print(f"FAIL: Component 1 -- No tables found (count: {len(doc.tables)})")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 8 and num_cols == 4:
                print(f"PASS: Component 1 -- Table has {num_rows} rows x {num_cols} cols (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 -- Expected 8x4 table, found {num_rows}x{num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Early exit if no table found
    if len(doc.tables) < 1:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 2: Header row contains correct text (0.25 points)
    # Expected headers: 'Item Description', 'Quantity', 'Unit Price', 'Total'
    try:
        expected_headers = ['Item Description', 'Quantity', 'Unit Price', 'Total']
        actual_headers = [cell.text.strip() for cell in table.rows[0].cells]
        matches = sum(1 for exp, act in zip(expected_headers, actual_headers) if exp == act)
        if len(actual_headers) >= 4 and matches == 4:
            print(f"PASS: Component 2 -- Header text matches: {actual_headers} (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 -- Expected {expected_headers}, found {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Header row text is bold (0.25 points)
    # Check that each header cell has at least one bold run with the header text
    try:
        bold_count = 0
        for j, cell in enumerate(table.rows[0].cells):
            cell_has_bold = False
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip() and run.font.bold:
                        cell_has_bold = True
                        break
                if cell_has_bold:
                    break
            if cell_has_bold:
                bold_count += 1

        if bold_count == 4:
            print(f"PASS: Component 3 -- All 4 header cells have bold text (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 -- {bold_count}/4 header cells have bold text")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Header row cells have bottom border (0.25 points)
    # Check for explicit bottom border on each header cell's tcPr/tcBorders
    try:
        border_count = 0
        for j, cell in enumerate(table.rows[0].cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                borders = tcPr.find(qn('w:tcBorders'))
                if borders is not None:
                    bottom = borders.find(qn('w:bottom'))
                    if bottom is not None:
                        val = bottom.get(qn('w:val'))
                        # 'nil' or 'none' means no border
                        if val and val not in ('nil', 'none'):
                            border_count += 1

        if border_count == 4:
            print(f"PASS: Component 4 -- All 4 header cells have bottom border (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 -- {border_count}/4 header cells have bottom border")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
