"""
Initial Setup: Insert a quarterly invoice table with bold header row and bottom border.
Task ID: writer_biz_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Company Header ---
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Meridian Consulting Group, LLC")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = subtitle.add_run("1247 Oakridge Boulevard, Suite 340\nPortland, OR 97205\nPhone: (503) 555-8192 | Fax: (503) 555-8193")
    r.font.size = Pt(10)
    r.font.name = "Calibri"

    # Spacing line
    doc.add_paragraph()

    # --- INVOICE Title ---
    inv_title = doc.add_paragraph()
    inv_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = inv_title.add_run("INVOICE")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Calibri"

    doc.add_paragraph()

    # --- Invoice Details ---
    details = doc.add_paragraph()
    details.paragraph_format.space_after = Pt(2)
    r = details.add_run("Invoice Number: ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r2 = details.add_run("INV-2025-0473")
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(2)
    r = date_para.add_run("Invoice Date: ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r2 = date_para.add_run("March 28, 2025")
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    due_para = doc.add_paragraph()
    due_para.paragraph_format.space_after = Pt(2)
    r = due_para.add_run("Due Date: ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r2 = due_para.add_run("April 27, 2025")
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    doc.add_paragraph()

    # --- Bill To ---
    bill_to = doc.add_paragraph()
    r = bill_to.add_run("Bill To:")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    client = doc.add_paragraph()
    r = client.add_run(
        "Brightwater Technologies, Inc.\n"
        "Attn: Jennifer Nakamura, Accounts Payable\n"
        "8901 Innovation Drive, Floor 12\n"
        "San Francisco, CA 94107"
    )
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    doc.add_paragraph()

    # --- Payment Terms ---
    terms = doc.add_paragraph()
    r = terms.add_run("Payment Terms: ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r2 = terms.add_run("Net 30")
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    doc.add_paragraph()

    # NO table here — the task is to insert one.

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
