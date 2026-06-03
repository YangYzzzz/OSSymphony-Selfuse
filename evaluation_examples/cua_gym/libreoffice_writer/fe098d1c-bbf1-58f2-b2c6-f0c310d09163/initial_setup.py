"""
Initial Setup: HR Update Document with Track Changes Task
Task ID: osworld_writer_comment_track_changes_004
Domain: libreoffice_writer

Creates a 4-paragraph HR update document with:
- "first quarter" present in paragraph 1
- paragraph 2 with 3+ sentences (third sentence to be deleted)
- "employees" in paragraph 3
- Track Changes off (initial state)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_comment_track_changes_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Quarterly HR Update — Internal Memo', level=1)

    # Paragraph 1 — contains "first quarter"
    p1 = doc.add_paragraph(
        'We are pleased to share the HR performance summary for the first quarter of 2025. '
        'Overall, recruitment efforts have exceeded our targets, and retention rates remain strong '
        'across all business units. Leadership is encouraged by the progress made during this period.'
    )

    # Paragraph 2 — contains 3 sentences; the third sentence is the one the agent will delete
    p2 = doc.add_paragraph(
        'The onboarding program has been redesigned to improve the experience for new hires. '
        'Feedback collected over the past three months indicates higher satisfaction scores compared to last year. '
        'Additionally, a new mentorship initiative is being piloted in the Engineering and Product divisions. '
        'We expect to roll out these improvements company-wide by the end of the second quarter.'
    )

    # Paragraph 3 — contains "employees"
    p3 = doc.add_paragraph(
        'All employees are encouraged to complete the updated compliance training by April 30th. '
        'The training covers recent regulatory changes and updated data privacy policies. '
        'HR will send reminders via email and the internal portal to ensure timely completion.'
    )

    # Paragraph 4 — closing
    p4 = doc.add_paragraph(
        'We appreciate the continued dedication and hard work demonstrated across the organization. '
        'If you have questions or concerns regarding any HR policies, please reach out to your HR Business Partner. '
        'Further updates will be communicated at the end of the month.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
