"""
Reward Script: Enable Track Changes, make edits, accept replacements, reject deletion
Task ID: osworld_writer_comment_track_changes_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): 'first quarter' replaced with 'Q1' in paragraph 1 (accepted change)
  Component 2 (0.5 pts): 'employees' replaced with 'team members' in paragraph 3 (accepted change)
  Note: Third sentence of paragraph 2 is expected to remain present (deletion rejected/restored).
        Its presence in both initial and golden means it cannot be scored, but its absence
        in golden would indicate incorrect task execution, captured as a diagnostic warning.
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_comment_track_changes_004'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.

    The task required:
      1. Enable Track Changes, then edit the document.
      2. Replace 'first quarter' with 'Q1' -> accept this change (incorporated into text).
      3. Delete third sentence of paragraph 2 -> reject this change (sentence restored in text).
      4. Change 'employees' to 'team members' -> accept this change (incorporated into text).

    Verifiable outcomes in the final file:
      - Paragraph 1 contains 'Q1' (not 'first quarter')
      - Paragraph 3 contains 'team members' (not 'employees')
      - Third sentence of paragraph 2 is still present (deletion was rejected)
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: Ensure expected document structure (5 paragraphs including heading)
    if len(doc.paragraphs) < 5:
        print(f"CRITICAL: Expected at least 5 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Paragraph index mapping:
    # Para 0: Heading
    # Para 1: First content paragraph (contains 'first quarter' / 'Q1')
    # Para 2: Second content paragraph (third sentence deletion)
    # Para 3: Third content paragraph (contains 'employees' / 'team members')
    # Para 4: Fourth content paragraph

    para1_text = doc.paragraphs[1].text
    para2_text = doc.paragraphs[2].text
    para3_text = doc.paragraphs[3].text

    # Diagnostic: Check third sentence of paragraph 2 is still present
    third_sentence_marker = "Additionally, a new mentorship initiative"
    if third_sentence_marker not in para2_text:
        print(f"DIAGNOSTIC WARNING: Third sentence of paragraph 2 appears to be missing. "
              f"Expected '{third_sentence_marker}' to be present (deletion should have been rejected). "
              f"Para 2 text: {repr(para2_text[:200])}")
    else:
        print("DIAGNOSTIC OK: Third sentence of paragraph 2 is present (deletion correctly rejected).")

    # Component 1: 'first quarter' replaced with 'Q1' in paragraph 1 (0.5 points)
    # Initial state has 'first quarter'; golden state should have 'Q1'.
    try:
        c1_q1_present = 'Q1' in para1_text
        c1_old_absent = 'first quarter' not in para1_text
        if c1_q1_present and c1_old_absent:
            print(f"PASS: Component 1 — 'Q1' found and 'first quarter' absent in paragraph 1 (0.5 pts). "
                  f"Snippet: {repr(para1_text[:100])}")
            total_score += 0.5
        elif not c1_old_absent:
            print(f"FAIL: Component 1 — 'first quarter' still present in paragraph 1 (not replaced). "
                  f"Snippet: {repr(para1_text[:100])}")
        else:
            print(f"FAIL: Component 1 — 'Q1' not found in paragraph 1. "
                  f"Snippet: {repr(para1_text[:100])}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'employees' replaced with 'team members' in paragraph 3 (0.5 points)
    # Initial state has 'employees'; golden state should have 'team members'.
    try:
        c2_new_present = 'team members' in para3_text
        c2_old_absent = 'employees' not in para3_text
        if c2_new_present and c2_old_absent:
            print(f"PASS: Component 2 — 'team members' found and 'employees' absent in paragraph 3 (0.5 pts). "
                  f"Snippet: {repr(para3_text[:100])}")
            total_score += 0.5
        elif not c2_old_absent:
            print(f"FAIL: Component 2 — 'employees' still present in paragraph 3 (not replaced). "
                  f"Snippet: {repr(para3_text[:100])}")
        else:
            print(f"FAIL: Component 2 — 'team members' not found in paragraph 3. "
                  f"Snippet: {repr(para3_text[:100])}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
