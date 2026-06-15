"""
Reward Script: Set footer margin distance to 1.0cm and add centered footer text
Task ID: writer_page_040
Domain: libreoffice_writer
Scoring:
  Component 1: Footer text is 'Confidential - Do Not Distribute' (0.5 pts)
  Component 2: Footer paragraph alignment is CENTER (0.2 pts)
  Component 3: Footer distance from page edge is ~1.0cm (0.3 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_040'
FILE_PATH = f'{WORKDIR}/confidential_memo.docx'

EXPECTED_FOOTER_TEXT = 'Confidential - Do Not Distribute'
EXPECTED_FOOTER_DISTANCE_CM = 1.0
TOLERANCE_CM = 0.05


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Checks:
    1. Footer contains the required text 'Confidential - Do Not Distribute'
    2. Footer paragraph alignment is CENTER
    3. Footer distance from bottom page edge is ~1.0cm (tolerance ±0.05cm)
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get first section (the document only has one section per context)
    try:
        section = doc.sections[0]
    except Exception as e:
        print(f"CRITICAL: Cannot access document section: {e}")
        print("REWARD: 0.0")
        return 0.0

    footer = section.footer

    # Component 1: Footer contains 'Confidential - Do Not Distribute' (0.5 points)
    try:
        footer_text = ''
        if footer.paragraphs:
            footer_text = footer.paragraphs[0].text.strip()

        if footer_text == EXPECTED_FOOTER_TEXT:
            print(f"PASS: Component 1 — Footer text matches '{EXPECTED_FOOTER_TEXT}' (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Expected footer text '{EXPECTED_FOOTER_TEXT}', found '{footer_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — Could not check footer text: {e}")

    # Component 2: Footer paragraph alignment is CENTER (0.2 points)
    try:
        if footer.paragraphs:
            para_alignment = footer.paragraphs[0].paragraph_format.alignment
            if para_alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                print(f"PASS: Component 2 — Footer alignment is CENTER (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 2 — Expected footer alignment CENTER, found {para_alignment}")
        else:
            print(f"FAIL: Component 2 — No footer paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 2 — Could not check footer alignment: {e}")

    # Component 3: Footer distance from bottom page edge is ~1.0cm (0.3 points)
    try:
        footer_distance_cm = section.footer_distance.cm
        if abs(footer_distance_cm - EXPECTED_FOOTER_DISTANCE_CM) <= TOLERANCE_CM:
            print(f"PASS: Component 3 — Footer distance is {footer_distance_cm:.4f}cm (expected ~{EXPECTED_FOOTER_DISTANCE_CM}cm, tolerance ±{TOLERANCE_CM}cm) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Footer distance is {footer_distance_cm:.4f}cm, expected ~{EXPECTED_FOOTER_DISTANCE_CM}cm (±{TOLERANCE_CM}cm)")
    except Exception as e:
        print(f"ERROR: Component 3 — Could not check footer distance: {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
