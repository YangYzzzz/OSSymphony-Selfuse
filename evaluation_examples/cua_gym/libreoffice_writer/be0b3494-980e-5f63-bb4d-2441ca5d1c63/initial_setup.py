"""
Initial Setup: Confidential memorandum - footer empty, default spacing
Task ID: writer_page_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_page_040'
# Context specifies file is at ~/Desktop/confidential_memo.docx
OUTPUT = f'{WORKDIR}/Desktop/confidential_memo.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Configure page: A4, portrait, margins all 2.54cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Footer: enabled but EMPTY (no text, default spacing)
    # Default footer_distance is typically 1.25cm; we leave it as default
    footer = section.footer
    footer.is_linked_to_previous = False
    # Ensure footer paragraph exists but has no text
    if footer.paragraphs:
        fp = footer.paragraphs[0]
        fp.clear()
    else:
        doc.add_paragraph()

    # Header: empty
    header = section.header
    header.is_linked_to_previous = False

    # ---- PAGE 1: Memorandum Header ----
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('CONFIDENTIAL MEMORANDUM')
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()  # blank line

    # Memo header fields
    memo_lines = [
        ('TO:', 'All Department Heads and Senior Management'),
        ('FROM:', 'Dr. Eleanor Whitfield, Chief Executive Officer'),
        ('DATE:', 'March 5, 2026'),
        ('RE:', 'Strategic Reorganization and Q1 Financial Overview'),
        ('CLASSIFICATION:', 'CONFIDENTIAL – Internal Use Only'),
    ]
    for label, value in memo_lines:
        p = doc.add_paragraph()
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_value = p.add_run('  ' + value)
        run_value.font.size = Pt(11)
    doc.add_paragraph()

    # Introduction paragraph
    intro = doc.add_paragraph()
    intro.add_run(
        'This memorandum is intended exclusively for the recipients listed above. '
        'Unauthorized disclosure, reproduction, or distribution of its contents is '
        'strictly prohibited and may result in disciplinary action or legal proceedings.'
    ).font.size = Pt(11)
    doc.add_paragraph()

    # Section 1
    h1 = doc.add_paragraph()
    h1_run = h1.add_run('1. Strategic Reorganization Update')
    h1_run.bold = True
    h1_run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.add_run(
        'Following the Board of Directors meeting held on February 20, 2026, the executive '
        'leadership team has approved a comprehensive reorganization of our global operations. '
        'The restructuring will consolidate three regional divisions into two integrated business '
        'units: North America & Europe (NAE) and Asia-Pacific & Emerging Markets (APEM).'
    ).font.size = Pt(11)
    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.add_run(
        'Ms. Rachel Okonkwo has been appointed as Executive Vice President of NAE, effective '
        'April 1, 2026. Mr. Thomas Liang will assume the role of EVP for APEM on the same date. '
        'Both appointments are subject to final board ratification on March 18, 2026. All '
        'department heads are requested to prepare transition briefings for their respective teams '
        'by March 25, 2026.'
    ).font.size = Pt(11)

    # Page break for page 2
    doc.add_page_break()

    # ---- PAGE 2: Financial Overview ----
    h2 = doc.add_paragraph()
    h2_run = h2.add_run('2. Q1 Financial Overview')
    h2_run.bold = True
    h2_run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.add_run(
        'Preliminary Q1 2026 financial results indicate strong performance across core business '
        'segments. Total consolidated revenue reached $847.3 million, representing a 12.4% '
        'increase year-over-year. Operating income improved to $124.6 million, with an EBITDA '
        'margin of 18.7%.'
    ).font.size = Pt(11)
    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.add_run(
        'Key highlights by division include: Engineering & Technology Services generated $312.4M '
        '(+15.2% YoY); Consumer Products & Retail recorded $289.7M (+9.8% YoY); Financial '
        'Services & Advisory contributed $245.2M (+12.1% YoY). Net earnings per share '
        '(diluted) of $2.34 surpassed analyst consensus estimates of $2.18 by 7.3%.'
    ).font.size = Pt(11)
    doc.add_paragraph()

    # Financial table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers_row = table.rows[0]
    headers_data = ['Division', 'Q1 2026 Revenue', 'YoY Growth']
    for i, h in enumerate(headers_data):
        cell = headers_row.cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    table_data = [
        ('Engineering & Technology Services', '$312.4M', '+15.2%'),
        ('Consumer Products & Retail', '$289.7M', '+9.8%'),
        ('Financial Services & Advisory', '$245.2M', '+12.1%'),
        ('Total Consolidated', '$847.3M', '+12.4%'),
    ]
    for i, (div, rev, growth) in enumerate(table_data, 1):
        row = table.rows[i]
        row.cells[0].paragraphs[0].add_run(div).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(rev).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(growth).font.size = Pt(10)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.add_run(
        'The CFO, Ms. Sandra Park, will present a detailed financial analysis during the '
        'all-hands leadership meeting on March 12, 2026. Attendance is mandatory for all '
        'senior vice presidents and above.'
    ).font.size = Pt(11)

    # Page break for page 3
    doc.add_page_break()

    # ---- PAGE 3: Compliance & Action Items ----
    h3 = doc.add_paragraph()
    h3_run = h3.add_run('3. Compliance and Governance Updates')
    h3_run.bold = True
    h3_run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.add_run(
        'The Legal and Compliance team has completed its annual review of corporate governance '
        'policies. Updated guidelines on data privacy, insider trading restrictions, and '
        'conflicts of interest will be distributed to all staff by March 15, 2026. Mandatory '
        'compliance training sessions are scheduled for March 20-24, 2026.'
    ).font.size = Pt(11)
    doc.add_paragraph()

    h4 = doc.add_paragraph()
    h4_run = h4.add_run('4. Action Items and Next Steps')
    h4_run.bold = True
    h4_run.font.size = Pt(13)

    action_items = [
        'Department Heads: Submit reorganization transition briefs by March 25, 2026.',
        'All VPs and above: Attend leadership meeting on March 12, 2026 at 09:00 HQ time.',
        'HR Team: Finalize new organizational charts and distribute by April 1, 2026.',
        'IT Security: Complete data migration risk assessment for APEM division by March 20.',
        'Finance Team: Prepare Q1 investor disclosure package per CFO guidance.',
        'Legal: Distribute updated compliance policies and schedule training sessions.',
    ]
    for item in action_items:
        li = doc.add_paragraph(style='List Bullet')
        li.add_run(item).font.size = Pt(11)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run(
        'Any questions or concerns regarding the matters outlined in this memorandum should '
        'be directed to the Office of the CEO at ceo-office@company.com. Please treat this '
        'document with the utmost confidentiality.'
    ).font.size = Pt(11)

    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.add_run('Dr. Eleanor Whitfield').bold = True
    sig_para.runs[-1].font.size = Pt(11)
    doc.add_paragraph().add_run('Chief Executive Officer').font.size = Pt(11)
    doc.add_paragraph().add_run('Date: March 5, 2026').font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
