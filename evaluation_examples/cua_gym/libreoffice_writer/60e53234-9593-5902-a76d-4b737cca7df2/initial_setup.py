"""
Initial Setup: Create a Writer document using only Default Page Style
Task ID: writer_bs_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_073'
# Create as .docx first, then convert to .odt for native page style support
DOCX_PATH = f'{WORKDIR}/{TASK_ID}.docx'
ODT_PATH = f'{WORKDIR}/{TASK_ID}.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup: A4 portrait, default margins ---
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Quarterly Business Review — Q1 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Solutions Group')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    run.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Prepared: March 28, 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'The first quarter of 2025 delivered strong results across all three '
        'business divisions. Revenue grew 14.2% year-over-year, driven primarily '
        'by the expansion of our enterprise SaaS platform and the successful '
        'launch of the Meridian Analytics suite in February. Client retention '
        'remained above target at 94.7%, while new client acquisition exceeded '
        'projections by 18%.'
    )
    doc.add_paragraph(
        'Operating expenses remained well-controlled, with total expenditure '
        'coming in 3.1% below budget. The engineering team completed the '
        'migration to cloud-native infrastructure ahead of schedule, which '
        'is expected to reduce hosting costs by approximately $230,000 annually '
        'starting in Q2.'
    )

    # --- Financial Highlights ---
    doc.add_heading('Financial Highlights', level=1)
    doc.add_paragraph(
        'Total revenue for Q1 2025 reached $4.87 million, compared to $4.26 '
        'million in Q1 2024. The breakdown by division is as follows:'
    )

    # Revenue table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Division', 'Q1 2025 Revenue', 'Q1 2024 Revenue', 'YoY Growth']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['Enterprise SaaS', '$2,340,000', '$1,980,000', '+18.2%'],
        ['Consulting Services', '$1,520,000', '$1,410,000', '+7.8%'],
        ['Analytics Platform', '$1,010,000', '$870,000', '+16.1%'],
        ['Total', '$4,870,000', '$4,260,000', '+14.3%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val
    # Bold total row
    for c in range(4):
        for run in table.cell(4, c).paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()  # spacer

    # --- Key Achievements ---
    doc.add_heading('Key Achievements', level=1)
    achievements = [
        'Launched Meridian Analytics v2.0 with predictive forecasting module',
        'Signed 12 new enterprise contracts (target was 10)',
        'Completed SOC 2 Type II certification ahead of schedule',
        'Reduced average customer onboarding time from 14 days to 8 days',
        'Hired 23 new team members across engineering and customer success',
    ]
    for item in achievements:
        doc.add_paragraph(item, style='List Bullet')

    # --- Departmental Updates ---
    doc.add_heading('Departmental Updates', level=1)

    doc.add_heading('Engineering', level=2)
    doc.add_paragraph(
        'The engineering division shipped 47 features and resolved 189 bugs '
        'during Q1. The team successfully migrated 92% of services to '
        'Kubernetes-based infrastructure, improving deployment frequency '
        'from bi-weekly to daily releases. Lead engineer Priya Ramanathan '
        'spearheaded the new CI/CD pipeline that reduced build times by 60%.'
    )

    doc.add_heading('Sales & Marketing', level=2)
    doc.add_paragraph(
        'The sales team closed $2.1 million in new annual recurring revenue. '
        'Marketing campaigns generated 3,400 qualified leads, a 22% increase '
        'over Q4 2024. The rebrand initiative led by Creative Director '
        'Sofia Martinez received positive feedback from 87% of surveyed clients.'
    )

    doc.add_heading('Customer Success', level=2)
    doc.add_paragraph(
        'Net Promoter Score improved to 72 (from 65 in Q4 2024). The support '
        'team maintained a 98.3% SLA compliance rate with an average first '
        'response time of 2.4 hours. Customer Success Manager Darius Okonkwo '
        'developed a new onboarding playbook that has been adopted company-wide.'
    )

    # --- Outlook ---
    doc.add_heading('Q2 2025 Outlook', level=1)
    doc.add_paragraph(
        'Looking ahead, we expect continued momentum across all divisions. '
        'Key priorities for Q2 include the launch of the Meridian Mobile '
        'application, expansion into the APAC market through our Singapore '
        'office, and the rollout of AI-assisted customer insights. Revenue '
        'guidance for Q2 is $5.1-5.4 million, representing 12-18% growth '
        'year-over-year.'
    )

    doc.save(DOCX_PATH)
    print(f'DOCX file created: {DOCX_PATH}')

    # Convert .docx to .odt using LibreOffice headless
    subprocess.run(['pkill', '-9', '-f', 'soffice'], capture_output=True)
    time.sleep(2)

    env = os.environ.copy()
    env['DISPLAY'] = ':0'
    result = subprocess.run(
        ['soffice', '--headless', '--convert-to', 'odt',
         '--outdir', WORKDIR, DOCX_PATH],
        capture_output=True, text=True, timeout=30, env=env
    )
    print(f'Conversion result: {result.stdout.strip()}')
    if result.stderr.strip():
        print(f'Conversion stderr: {result.stderr.strip()}')

    time.sleep(2)

    # Remove the .docx to avoid confusion
    if os.path.exists(ODT_PATH):
        os.remove(DOCX_PATH)
        print(f'Initial file ready: {ODT_PATH}')
    else:
        print(f'WARNING: ODT file not created, keeping DOCX')

    # GUI-ready startup - open the .odt file
    target = ODT_PATH if os.path.exists(ODT_PATH) else DOCX_PATH
    # Kill any leftover soffice from conversion
    subprocess.run(['pkill', '-9', '-f', 'soffice'], capture_output=True)
    time.sleep(2)
    launch_gui(f'libreoffice --writer "{target}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
