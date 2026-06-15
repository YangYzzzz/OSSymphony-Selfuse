"""
FINAL REWARD SCRIPT - SUCCESS
Task: I’ve got a LibreOffice Writer document where the very first paragraph lists a bunch of figures, e.g. 245, 1 400, 37, etc. For compliance I need every one of those numbers in paragraph 1 to carry a double underline—no letters, no punctuation, just the digits themselves. What’s the quickest way to select only those numbers in that single paragraph and apply the built-in “Double” underline style?
Generated: 2025-09-10 12:29:52
Status: success
Model: azure-o3
Total Steps: 1
"""

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
import os


def is_double_underlined(run):
    """Check if a run has a *double* underline applied.
    Handles both high-level python-docx API and direct XML inspection.
    """
    try:
        val = run.font.underline
        # python-docx maps Word underline values to WD_UNDERLINE enums or strings
        if val == WD_UNDERLINE.DOUBLE or (isinstance(val, str) and str(val).lower() == "double"):
            return True
    except Exception:
        pass

    # Fallback: inspect the run's XML for <w:u w:val="double"/>
    u_el = run._element.find(".//w:u", run._element.nsmap)
    if u_el is not None:
        val_attr = u_el.get(qn("w:val")) or u_el.get("w:val")
        if val_attr and val_attr.lower() == "double":
            return True
    return False


def verify_double_underline_digits(file_path):
    """Reward-function to verify that *only* digits in the first paragraph
    are double-underlined, and that *all* such digits carry the formatting.
    Progressive scoring:
      • 60% weight – proportion of digit characters correctly double-underlined
      • 40% weight – proportion of non-digit characters *not* double-underlined
    Returns a float in [0.0, 1.0].
    """
    print(f"Verifying document: {file_path}")
    if not os.path.exists(file_path):
        print("✗ File not found")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to load DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    if not doc.paragraphs:
        print("✗ Document contains no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    first_para = doc.paragraphs[0]
    para_text = first_para.text
    print(f"First paragraph text: '{para_text}'")

    # Count digit positions
    digit_positions = [i for i, ch in enumerate(para_text) if ch.isdigit()]
    total_digits = len(digit_positions)
    if total_digits == 0:
        print("✗ No digits present – nothing to verify")
        print("REWARD: 0.0")
        return 0.0
    print(f"Total digit characters detected: {total_digits}")

    # Iterate through runs to evaluate underline formatting per character
    digit_correct, digit_incorrect = 0, 0
    nondigit_total, nondigit_incorrect = 0, 0

    for run in first_para.runs:
        if not run.text:
            continue
        run_double = is_double_underlined(run)
        for char in run.text:
            if char.isdigit():
                if run_double:
                    digit_correct += 1
                else:
                    digit_incorrect += 1
            else:
                nondigit_total += 1
                if run_double:
                    nondigit_incorrect += 1

    print(f"Digits correctly double-underlined: {digit_correct}/{total_digits}")
    print(f"Digits missing double underline: {digit_incorrect}")
    print(f"Non-digit chars wrongly double-underlined: {nondigit_incorrect}/{nondigit_total}")

    # Progressive scoring
    score = 0.0
    # 60% for digits correctness
    score += (digit_correct / total_digits) * 0.6
    # 40% for non-digits *not* being double-underlined
    if nondigit_total:
        score += ((nondigit_total - nondigit_incorrect) / nondigit_total) * 0.4
    else:
        score += 0.4  # no non-digit characters means this requirement is trivially met

    score = round(min(score, 1.0), 4)
    print(f"Computed score: {score}")
    print(f"REWARD: {score}")
    return score


if __name__ == "__main__":
    # Path provided by the task context
    file_path = "/home/user/ive_got_a_libreoffice_writer_document_where_the_very_first_paragraph_lists_a_bunch_of_figures_eg_245.docx"
    verify_double_underline_digits(file_path)
