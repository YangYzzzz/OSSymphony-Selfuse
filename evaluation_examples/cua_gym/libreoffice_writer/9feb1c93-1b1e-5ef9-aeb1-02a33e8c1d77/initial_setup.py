"""
Initial Setup: Mail merge field placeholders in client letter template
Task ID: writer_biz_057
Domain: libreoffice_writer

Creates a professional client letter template with placeholder text
[Client Name], [Company], and [Address]. Also creates a data source
spreadsheet with Name, Company, and Address columns.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_057'
OUTPUT_DOCX = f'{WORKDIR}/{TASK_ID}.docx'
OUTPUT_XLSX = f'{WORKDIR}/client_data.xlsx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_data_source():
    """Create a spreadsheet data source with client information."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Clients"

    headers = ['Name', 'Company', 'Address']
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)

    clients = [
        ['Sarah Chen', 'Meridian Technologies Inc.', '1420 Innovation Drive, Suite 300, San Jose, CA 95134'],
        ['Marcus Johnson', 'Apex Financial Group', '875 Market Street, Floor 12, New York, NY 10013'],
        ['Elena Rodriguez', 'Greenfield Consulting Partners', '2200 Lakeshore Boulevard, Chicago, IL 60614'],
        ['David Park', 'Summit Healthcare Solutions', '9500 Gilman Drive, La Jolla, CA 92093'],
        ['Amara Okafor', 'BrightPath Education Corp.', '300 Congress Avenue, Austin, TX 78701'],
    ]

    for r, row_data in enumerate(clients, 2):
        for c, val in enumerate(row_data, 1):
            ws.cell(row=r, column=c, value=val)

    # Auto-size columns
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 35
    ws.column_dimensions['C'].width = 55

    wb.save(OUTPUT_XLSX)
    print(f'Data source created: {OUTPUT_XLSX}')


def create_letter_template():
    """Create a client letter template with placeholder text."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Company letterhead
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run('Cascade Business Solutions')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub_heading = doc.add_paragraph()
    sub_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub_heading.add_run('4500 Executive Park Drive, Suite 220, Portland, OR 97201')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = contact.add_run('Tel: (503) 555-0184  |  info@cascadebusiness.com  |  www.cascadebusiness.com')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Horizontal line (via border on empty paragraph)
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(6)
    line_para.paragraph_format.space_after = Pt(12)
    pPr = line_para._element.get_or_add_pPr()
    from docx.oxml.ns import qn
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '1A3C6E',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Date
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(12)
    run = date_para.add_run('March 28, 2026')
    run.font.size = Pt(11)

    # Blank line
    doc.add_paragraph()

    # Recipient block with placeholders
    recipient = doc.add_paragraph()
    recipient.paragraph_format.space_after = Pt(0)
    run = recipient.add_run('[Client Name]')
    run.font.size = Pt(11)

    company_line = doc.add_paragraph()
    company_line.paragraph_format.space_before = Pt(0)
    company_line.paragraph_format.space_after = Pt(0)
    run = company_line.add_run('[Company]')
    run.font.size = Pt(11)

    address_line = doc.add_paragraph()
    address_line.paragraph_format.space_before = Pt(0)
    address_line.paragraph_format.space_after = Pt(12)
    run = address_line.add_run('[Address]')
    run.font.size = Pt(11)

    # Salutation
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(6)
    run = salutation.add_run('Dear [Client Name],')
    run.font.size = Pt(11)

    # Body paragraphs
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(6)
    run = body1.add_run(
        'Thank you for choosing Cascade Business Solutions as your strategic partner. '
        'We are writing to confirm the details of our upcoming quarterly business review '
        'meeting scheduled for April 15, 2026.'
    )
    run.font.size = Pt(11)

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(6)
    run = body2.add_run(
        'During this session, we will review the performance metrics from Q1 2026, '
        'discuss the implementation timeline for Phase 2 of the digital transformation '
        'initiative, and outline the strategic objectives for the remainder of the fiscal year. '
        'Our analytics team has prepared a comprehensive report that highlights key areas of '
        'growth and identifies opportunities for operational improvement.'
    )
    run.font.size = Pt(11)

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(6)
    run = body3.add_run(
        'Please confirm your availability at your earliest convenience. If you require '
        'any additional materials or have specific topics you would like to add to the agenda, '
        'do not hesitate to contact our project coordinator, Rebecca Torres, at '
        'rtorres@cascadebusiness.com or (503) 555-0192.'
    )
    run.font.size = Pt(11)

    body4 = doc.add_paragraph()
    body4.paragraph_format.space_after = Pt(12)
    run = body4.add_run(
        'We look forward to continuing our productive collaboration with [Company] '
        'and are committed to delivering measurable results that align with your '
        'organizational goals.'
    )
    run.font.size = Pt(11)

    # Closing
    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(24)
    run = closing.add_run('Sincerely,')
    run.font.size = Pt(11)

    # Signature
    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_after = Pt(0)
    run = sig_name.add_run('Jonathan R. Mitchell')
    run.bold = True
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_before = Pt(0)
    sig_title.paragraph_format.space_after = Pt(0)
    run = sig_title.add_run('Senior Partner, Client Relations')
    run.font.size = Pt(11)

    sig_company = doc.add_paragraph()
    sig_company.paragraph_format.space_before = Pt(0)
    run = sig_company.add_run('Cascade Business Solutions')
    run.font.size = Pt(11)

    doc.save(OUTPUT_DOCX)
    print(f'Letter template created: {OUTPUT_DOCX}')


def main():
    create_data_source()
    create_letter_template()

    # Launch LibreOffice Writer with the template
    launch_gui(f'libreoffice --writer "{OUTPUT_DOCX}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
