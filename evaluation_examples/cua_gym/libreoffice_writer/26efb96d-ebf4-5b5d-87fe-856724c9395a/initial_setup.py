"""
Initial Setup: Configure paragraph 3 keep-with-next in a 12-paragraph report
Task ID: wrpara_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_016'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins for a standard report
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Paragraph 1: Report Title ---
    title = doc.add_heading('Quarterly Performance Review — Q4 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Paragraph 2: Executive Summary ---
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    run = p2.add_run(
        'This report provides a comprehensive overview of the company\'s performance '
        'during the fourth quarter of fiscal year 2025. The analysis covers revenue growth '
        'across all divisions, operational efficiency metrics, customer acquisition trends, '
        'and workforce development outcomes. Key findings indicate a 14.2% year-over-year '
        'increase in consolidated revenue, driven primarily by strong growth in the '
        'enterprise solutions segment and expanding market share in the Asia-Pacific region. '
        'Operating margins improved by 2.3 percentage points compared to Q3 2025, reflecting '
        'ongoing cost optimization initiatives and favorable product mix shifts. Employee '
        'retention rates remained above industry benchmarks at 91.7%, supported by enhanced '
        'benefits programs introduced in mid-2025. The board of directors has reviewed these '
        'results and approved the strategic expansion plan for fiscal year 2026, which '
        'includes targeted investments in artificial intelligence capabilities and a new '
        'regional headquarters in Singapore.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Paragraph 3: Section Heading "Analysis Results" ---
    # This is the target paragraph — it must NOT have keep_with_next set initially.
    heading3 = doc.add_heading('Analysis Results', level=1)

    # --- Paragraph 4: Revenue Analysis ---
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(10)
    run = p4.add_run(
        'Total revenue for Q4 2025 reached $127.8 million, exceeding the forecast of '
        '$118.5 million by 7.8%. The enterprise solutions division contributed $54.3 million, '
        'representing a 22.1% increase over the same period last year. Consumer products '
        'generated $41.2 million in revenue, while professional services accounted for the '
        'remaining $32.3 million. The Asia-Pacific region was the strongest performer, with '
        'revenue growth of 31.4% driven by new partnerships with regional distributors in '
        'Japan, South Korea, and Australia. The European market showed moderate growth of '
        '8.7%, while North American revenue grew by 11.3% despite increased competitive '
        'pressure from emerging market entrants.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Paragraph 5: Operational Metrics ---
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(10)
    run = p5.add_run(
        'Operating efficiency improved significantly during Q4 2025. The cost-to-revenue '
        'ratio decreased from 72.4% in Q3 to 70.1%, reflecting successful implementation '
        'of automated workflow systems across the manufacturing and logistics departments. '
        'Average order fulfillment time was reduced from 4.2 business days to 3.1 business '
        'days, a 26.2% improvement. Customer support ticket resolution time averaged '
        '18.4 hours, down from 24.7 hours in the previous quarter. These improvements '
        'were achieved while simultaneously expanding headcount by 4.8%, indicating '
        'genuine productivity gains rather than simple cost-cutting measures.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Paragraph 6: Another Section Heading ---
    doc.add_heading('Customer Acquisition and Retention', level=1)

    # --- Paragraph 7: Customer metrics ---
    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(10)
    run = p7.add_run(
        'New customer acquisition in Q4 2025 totaled 3,847 accounts, a 16.9% increase '
        'over Q3 2025 and a 28.3% increase year-over-year. The average customer acquisition '
        'cost (CAC) was $342 per account, down from $387 in Q3, reflecting improved '
        'targeting algorithms in digital marketing campaigns. Customer lifetime value (CLV) '
        'projections increased to $12,450 per enterprise account and $2,180 per SMB account. '
        'Net Promoter Score (NPS) improved from 47 to 52 during the quarter, placing the '
        'company in the top quartile of its industry peer group.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Paragraph 8: Churn analysis ---
    p8 = doc.add_paragraph()
    p8.paragraph_format.space_after = Pt(10)
    run = p8.add_run(
        'Monthly churn rate stabilized at 1.8%, below the industry average of 2.4%. '
        'The implementation of proactive engagement workflows identified 412 at-risk '
        'accounts during Q4, of which 78.6% were successfully retained through targeted '
        'outreach and customized retention offers. Annual recurring revenue (ARR) from '
        'existing customers grew by 9.2%, driven by upsell conversions in the mid-market '
        'segment. The customer success team expanded from 28 to 35 members to support '
        'the growing account base, with each representative managing an average portfolio '
        'of $2.8 million in ARR.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Paragraph 9: Another Section Heading ---
    doc.add_heading('Workforce Development', level=1)

    # --- Paragraph 10: Employee metrics ---
    p10 = doc.add_paragraph()
    p10.paragraph_format.space_after = Pt(10)
    run = p10.add_run(
        'Total headcount at the end of Q4 2025 was 1,247 full-time employees, an increase '
        'of 57 positions from Q3. The engineering division added 31 new hires, primarily in '
        'machine learning and cloud infrastructure roles. Average time-to-fill for open '
        'positions decreased from 42 days to 35 days, supported by enhanced employer '
        'branding initiatives and a refreshed referral bonus program. Training investment '
        'totaled $1.4 million during the quarter, with 89% of employees completing at least '
        'one professional development course.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Paragraph 11: Diversity and inclusion ---
    p11 = doc.add_paragraph()
    p11.paragraph_format.space_after = Pt(10)
    run = p11.add_run(
        'Diversity hiring targets were exceeded in Q4, with 43.8% of new hires from '
        'underrepresented groups, surpassing the 40% target. Gender representation in '
        'leadership positions improved from 34.2% to 36.7% women. The employee engagement '
        'survey conducted in November 2025 returned a participation rate of 87.3% and an '
        'overall engagement score of 4.1 out of 5.0, both representing improvements over '
        'the mid-year survey. Key areas identified for improvement include cross-departmental '
        'communication and career progression transparency.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Paragraph 12: Conclusion ---
    p12 = doc.add_paragraph()
    p12.paragraph_format.space_after = Pt(10)
    run = p12.add_run(
        'In conclusion, Q4 2025 represents a strong finish to the fiscal year, with '
        'performance exceeding expectations across all major metrics. The combination of '
        'revenue growth, operational improvements, and workforce investments positions the '
        'company favorably for the strategic expansion planned in fiscal year 2026. '
        'Management recommends continued investment in AI-driven product capabilities, '
        'further expansion into Asia-Pacific markets, and sustained focus on employee '
        'development and retention as priorities for the coming year. The next quarterly '
        'review will be presented to the board on April 15, 2026.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
