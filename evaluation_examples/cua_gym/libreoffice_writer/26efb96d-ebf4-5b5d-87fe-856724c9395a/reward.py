"""
Reward Script: Configure paragraph 3 ("Analysis Results" heading) to keep with next paragraph
Task ID: wrpara_016
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): Paragraph 3 (index 2) has keepNext enabled
  Component 2 (0.2): Paragraph 3 text is still "Analysis Results" (not corrupted)
  Component 3 (0.2): keepNext is specifically on paragraph index 2 and NOT already on other non-heading paragraphs
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'wrpara_016'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def check_keep_next(para):
    """Check if a paragraph has w:keepNext enabled in its XML properties."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return False
    kn = pPr.find(qn('w:keepNext'))
    if kn is None:
        return False
    # If val attribute is absent, it defaults to true; if present, check it's not "false"/"0"
    val = kn.get(qn('w:val'))
    if val is None:
        return True
    return val.lower() not in ('false', '0')


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least 3 paragraphs
    if len(doc.paragraphs) < 3:
        print(f"FAIL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 3")
        print("REWARD: 0.0")
        return 0.0

    para_2 = doc.paragraphs[2]  # paragraph 3 (1-indexed) = index 2

    # Component 1: Paragraph 3 (index 2) has keepNext enabled (0.6 points)
    # This is the PRIMARY task requirement — "Keep with next paragraph" must be on
    try:
        has_keep_next = check_keep_next(para_2)
        if has_keep_next:
            print(f"PASS: Component 1 — Paragraph 3 has keepNext enabled (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 — Paragraph 3 does NOT have keepNext enabled")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraph 3 text content is preserved as "Analysis Results" (0.2 points)
    # The task should only change formatting, not content. We verify the heading text
    # is intact. This component FAILS on initial because we combine it with keepNext.
    # Actually, the text is the same on both envs, so we must tie it to the change.
    # Revised: Award points only if keepNext is enabled AND text is preserved.
    try:
        text = para_2.text.strip()
        if has_keep_next and text == "Analysis Results":
            print(f"PASS: Component 2 — keepNext enabled AND heading text preserved as 'Analysis Results' (0.2 pts)")
            total_score += 0.2
        elif not has_keep_next:
            print(f"FAIL: Component 2 — keepNext not enabled (prerequisite failed)")
        else:
            print(f"FAIL: Component 2 — Heading text changed to '{text}', expected 'Analysis Results'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: keepNext on para 3 AND no spurious keepNext added to normal body paragraphs (0.2 points)
    # We check that the change is specifically targeted: keepNext is on the heading (index 2)
    # but NOT added to non-heading body paragraphs (indices 1, 3, 4, 6, 7, 9, 10, 11)
    try:
        non_heading_indices = [1, 3, 4, 6, 7, 9, 10, 11]
        spurious_keep_next = []
        for idx in non_heading_indices:
            if idx < len(doc.paragraphs):
                if check_keep_next(doc.paragraphs[idx]):
                    spurious_keep_next.append(idx)

        if has_keep_next and len(spurious_keep_next) == 0:
            print(f"PASS: Component 3 — keepNext correctly targeted to paragraph 3 only, no spurious changes (0.2 pts)")
            total_score += 0.2
        elif not has_keep_next:
            print(f"FAIL: Component 3 — keepNext not enabled on paragraph 3 (prerequisite failed)")
        else:
            print(f"FAIL: Component 3 — Spurious keepNext found on paragraph indices: {spurious_keep_next}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
