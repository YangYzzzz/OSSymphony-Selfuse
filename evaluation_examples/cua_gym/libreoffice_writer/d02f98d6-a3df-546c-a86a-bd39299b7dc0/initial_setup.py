"""
Initial Setup: Grant application document with 8 tracked changes
Task ID: writer_struct_045
Domain: libreoffice_writer

Creates grant_application.docx on ~/Desktop/ with:
- 5-page grant proposal document
- 8 tracked changes total (including one that replaces '$10,000' with '$15,000' on page 2)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'grant_application'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_del_run(author, date, run_text, rpr_xml=None):
    """Create a w:del element containing a w:delText run."""
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), '1')
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date)

    del_run = OxmlElement('w:r')
    if rpr_xml is not None:
        del_run.append(copy.deepcopy(rpr_xml))
    del_text = OxmlElement('w:delText')
    del_text.set(qn('xml:space'), 'preserve')
    del_text.text = run_text
    del_run.append(del_text)
    del_elem.append(del_run)
    return del_elem


def make_ins_run(author, date, run_text, rpr_xml=None):
    """Create a w:ins element containing a w:r run."""
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:id'), '2')
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), date)

    ins_run = OxmlElement('w:r')
    if rpr_xml is not None:
        ins_run.append(copy.deepcopy(rpr_xml))
    ins_text = OxmlElement('w:t')
    ins_text.set(qn('xml:space'), 'preserve')
    ins_text.text = run_text
    ins_run.append(ins_text)
    ins_elem.append(ins_run)
    return ins_elem


def add_tracked_replacement(para, before_text, del_text_str, ins_text_str, after_text,
                             author, date, change_id, rpr_xml=None):
    """
    Replace para content with: [before_text][del del_text_str][ins ins_text_str][after_text]
    This simulates a tracked change that replaces del_text_str with ins_text_str.
    """
    # Clear existing runs
    p_elem = para._p
    # Remove all r, ins, del children
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'ins', 'del', 'hyperlink'):
            p_elem.remove(child)

    def add_normal_run(text):
        r = OxmlElement('w:r')
        if rpr_xml is not None:
            r.append(copy.deepcopy(rpr_xml))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p_elem.append(r)

    if before_text:
        add_normal_run(before_text)

    # del element
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), str(change_id))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date)
    del_run = OxmlElement('w:r')
    if rpr_xml is not None:
        del_run.append(copy.deepcopy(rpr_xml))
    del_t = OxmlElement('w:delText')
    del_t.set(qn('xml:space'), 'preserve')
    del_t.text = del_text_str
    del_run.append(del_t)
    del_elem.append(del_run)
    p_elem.append(del_elem)

    # ins element
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:id'), str(change_id + 1))
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), date)
    ins_run = OxmlElement('w:r')
    if rpr_xml is not None:
        ins_run.append(copy.deepcopy(rpr_xml))
    ins_t = OxmlElement('w:t')
    ins_t.set(qn('xml:space'), 'preserve')
    ins_t.text = ins_text_str
    ins_run.append(ins_t)
    ins_elem.append(ins_run)
    p_elem.append(ins_elem)

    if after_text:
        add_normal_run(after_text)


def add_ins_paragraph(doc, text, author, date, change_id, style=None):
    """Add a paragraph whose entire content is a tracked insertion."""
    if style:
        para = doc.add_paragraph(style=style)
    else:
        para = doc.add_paragraph()
    p_elem = para._p
    # Remove default run
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'ins', 'del'):
            p_elem.remove(child)

    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:id'), str(change_id))
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), date)
    ins_run = OxmlElement('w:r')
    ins_t = OxmlElement('w:t')
    ins_t.set(qn('xml:space'), 'preserve')
    ins_t.text = text
    ins_run.append(ins_t)
    ins_elem.append(ins_run)
    p_elem.append(ins_elem)
    return para


def add_del_paragraph_text(para, text, author, date, change_id):
    """Mark an existing paragraph's text as deleted (tracked deletion)."""
    p_elem = para._p
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'ins', 'del'):
            p_elem.remove(child)

    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), str(change_id))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date)
    del_run = OxmlElement('w:r')
    del_t = OxmlElement('w:delText')
    del_t.set(qn('xml:space'), 'preserve')
    del_t.text = text
    del_run.append(del_t)
    del_elem.append(del_run)
    p_elem.append(del_elem)


def add_format_change_run(para, text, author, date, change_id, old_bold=None, new_bold=True):
    """Add a run with a formatting tracked change (rPrChange)."""
    p_elem = para._p
    r = OxmlElement('w:r')

    rpr = OxmlElement('w:rPr')
    if new_bold:
        b = OxmlElement('w:b')
        rpr.append(b)

    # rPrChange records old formatting
    rpr_change = OxmlElement('w:rPrChange')
    rpr_change.set(qn('w:id'), str(change_id))
    rpr_change.set(qn('w:author'), author)
    rpr_change.set(qn('w:date'), date)
    old_rpr = OxmlElement('w:rPr')
    if old_bold:
        old_b = OxmlElement('w:b')
        old_rpr.append(old_b)
    rpr_change.append(old_rpr)
    rpr.append(rpr_change)
    r.append(rpr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_elem.append(r)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---- PAGE 1: Title and Introduction ----
    title_para = doc.add_heading('Community Health Initiative Grant Application', level=0)
    title_para.alignment = 1  # CENTER

    # Applicant info
    doc.add_paragraph('Applicant Organization: Riverside Community Health Foundation')
    doc.add_paragraph('Project Title: Expanding Access to Preventive Healthcare in Underserved Communities')
    doc.add_paragraph('Funding Period: January 2026 – December 2028')
    doc.add_paragraph('Total Funding Requested: $245,000')

    doc.add_heading('1. Executive Summary', level=1)
    p_exec = doc.add_paragraph(
        'The Riverside Community Health Foundation requests funding to expand preventive healthcare '
        'services to underserved populations in the greater metropolitan area. This three-year initiative '
        'will establish two mobile health clinics, train 15 community health workers, and provide '
        'free screenings to an estimated 3,500 individuals annually.'
    )

    # Tracked change 1: insertion of a sentence (new content added by reviewer)
    add_ins_paragraph(
        doc,
        'The project builds on our successful 2023–2025 pilot program, which served over 800 residents.',
        author='Dr. Patricia Nguyen',
        date='2026-02-10T09:15:00Z',
        change_id=1
    )

    doc.add_heading('2. Organizational Background', level=1)
    doc.add_paragraph(
        'Founded in 2010, the Riverside Community Health Foundation (RCHF) has a 15-year history of '
        'delivering health education, prevention programs, and direct service to low-income families. '
        'Our annual budget of $1.2 million supports six full-time staff and a network of over 40 volunteers.'
    )

    # Tracked change 2: text replacement in organizational background
    bg_para = doc.add_paragraph()
    add_tracked_replacement(
        bg_para,
        before_text='RCHF holds accreditation from the National Community Health Association and maintains ',
        del_text_str='a strong',
        ins_text_str='an excellent',
        after_text=' track record of fiscal responsibility and program outcomes.',
        author='James Whitfield',
        date='2026-02-12T14:30:00Z',
        change_id=3
    )

    doc.add_page_break()

    # ---- PAGE 2: Budget Justification ----
    doc.add_heading('3. Budget Justification', level=1)
    doc.add_paragraph(
        'The proposed budget of $245,000 over three years reflects the true cost of delivering high-quality '
        'community health services. Each line item has been carefully considered and documented.'
    )

    doc.add_heading('3.1 Personnel Costs', level=2)
    doc.add_paragraph(
        'Personnel costs account for 62% of the total budget ($151,900). This includes '
        'two full-time Community Health Workers (CHWs) at $42,000 each annually, a part-time '
        'Program Coordinator at $28,000 annually, and fringe benefits at 18%.'
    )

    doc.add_heading('3.2 Equipment and Supplies', level=2)

    # PAGE 2 KEY PARAGRAPH: tracked change replaces '$10,000' with '$15,000'
    equip_para = doc.add_paragraph()
    add_tracked_replacement(
        equip_para,
        before_text='Medical equipment and diagnostic supplies are budgeted at ',
        del_text_str='$10,000',
        ins_text_str='$15,000',
        after_text=' per year, covering blood pressure monitors, glucose testing kits, '
                   'vision screening tools, and basic wound care materials. This estimate is '
                   'based on current supplier quotes and historical usage data from our pilot program.',
        author='Dr. Patricia Nguyen',
        date='2026-02-14T11:45:00Z',
        change_id=5
    )

    doc.add_heading('3.3 Travel and Transportation', level=2)
    doc.add_paragraph(
        'Mobile clinic fuel and maintenance costs are estimated at $8,400 annually, based on '
        '200 service days per year with an average round-trip distance of 45 miles at $0.67/mile '
        'plus routine maintenance of $1,200/year per vehicle.'
    )

    doc.add_heading('3.4 Indirect Costs', level=2)

    # Tracked change 3: formatting change (bold added to a keyword)
    indirect_para = doc.add_paragraph('Indirect costs are calculated at ')
    add_format_change_run(
        indirect_para,
        '10% of direct costs',
        author='James Whitfield',
        date='2026-02-15T10:00:00Z',
        change_id=7,
        old_bold=False,
        new_bold=True
    )
    indirect_para._p.append(
        _make_normal_run_elem(' ($22,500 total), consistent with our federally negotiated indirect cost rate agreement.')
    )

    doc.add_page_break()

    # ---- PAGE 3: Program Design ----
    doc.add_heading('4. Program Design and Implementation', level=1)
    doc.add_paragraph(
        'The program will operate through a hub-and-spoke model, with a central coordination office '
        'and two mobile units serving six target neighborhoods identified through our 2024 Community '
        'Health Needs Assessment.'
    )

    doc.add_heading('4.1 Mobile Health Clinics', level=2)

    # Tracked change 4: deletion of a sentence
    mobile_del_para = doc.add_paragraph()
    add_del_paragraph_text(
        mobile_del_para,
        'Each mobile unit will be staffed by one CHW and one licensed practical nurse during service days.',
        author='Dr. Patricia Nguyen',
        date='2026-02-16T08:30:00Z',
        change_id=8
    )

    doc.add_paragraph(
        'Mobile units will visit each neighborhood twice monthly, offering screenings for hypertension, '
        'diabetes, vision impairment, and oral health. Referrals will be coordinated with partnering '
        'clinics including Riverside General Hospital and Community Care Clinic.'
    )

    doc.add_heading('4.2 Community Health Worker Training', level=2)
    doc.add_paragraph(
        'All CHWs will complete a 120-hour certification program accredited by the State Department '
        'of Health, followed by 40 hours of program-specific training covering motivational interviewing, '
        'cultural competency, and electronic health record documentation.'
    )

    # Tracked change 5: text replacement in training section
    train_para = doc.add_paragraph()
    add_tracked_replacement(
        train_para,
        before_text='CHWs will receive ongoing supervision from the Program Coordinator and participate in ',
        del_text_str='monthly',
        ins_text_str='bi-weekly',
        after_text=' peer learning sessions to share best practices and address challenges.',
        author='James Whitfield',
        date='2026-02-17T13:15:00Z',
        change_id=9
    )

    doc.add_page_break()

    # ---- PAGE 4: Evaluation ----
    doc.add_heading('5. Evaluation Plan', level=1)
    doc.add_paragraph(
        'Program effectiveness will be measured through a mixed-methods evaluation framework '
        'developed in partnership with the University of Riverside School of Public Health.'
    )

    doc.add_heading('5.1 Process Measures', level=2)
    doc.add_paragraph(
        'Process measures include: number of individuals screened per quarter, percentage of '
        'referrals completed within 30 days, CHW training completion rates, and mobile unit '
        'service day adherence.'
    )

    # Tracked change 6: insertion of a new bullet point
    add_ins_paragraph(
        doc,
        'Client satisfaction surveys will be administered quarterly using validated instruments.',
        author='Dr. Patricia Nguyen',
        date='2026-02-18T09:45:00Z',
        change_id=11
    )

    doc.add_heading('5.2 Outcome Measures', level=2)
    doc.add_paragraph(
        'Primary outcomes include: reduction in uncontrolled hypertension rates (target: 15% reduction '
        'by year 3), increased diabetes screening rates (target: 25% increase), and improved early '
        'detection of vision problems (target: 200 referrals annually).'
    )

    # Tracked change 7: text replacement in outcome section
    outcome_para = doc.add_paragraph()
    add_tracked_replacement(
        outcome_para,
        before_text='Secondary outcomes include increased health literacy scores as measured by the ',
        del_text_str='Rapid Estimate of Adult Literacy in Medicine (REALM)',
        ins_text_str='Newest Vital Sign (NVS) health literacy assessment',
        after_text=' administered at baseline and annually thereafter.',
        author='James Whitfield',
        date='2026-02-19T11:00:00Z',
        change_id=12
    )

    doc.add_page_break()

    # ---- PAGE 5: Sustainability and Conclusion ----
    doc.add_heading('6. Sustainability Plan', level=1)
    doc.add_paragraph(
        'RCHF is committed to sustaining the program beyond the grant period. Revenue diversification '
        'strategies include Medicaid billing for eligible services (projected to generate $35,000 annually '
        'by year 2), corporate wellness partnerships with three identified local employers, and inclusion '
        'in the city\'s annual community benefit budget.'
    )

    # Tracked change 8: deletion in sustainability section
    sustain_del_para = doc.add_paragraph()
    add_del_paragraph_text(
        sustain_del_para,
        'A detailed sustainability analysis conducted by our financial consultant projects 85% cost recovery by year 3.',
        author='Dr. Patricia Nguyen',
        date='2026-02-20T15:30:00Z',
        change_id=14
    )

    doc.add_paragraph(
        'We have already secured letters of commitment from two corporate partners and are in advanced '
        'negotiations with a third. These partnerships, combined with Medicaid revenues, will cover '
        'approximately 70% of annual operating costs by the end of the grant period.'
    )

    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        'The Riverside Community Health Foundation\'s proposed initiative represents a cost-effective, '
        'evidence-based approach to reducing health disparities in our community. With 15 years of '
        'proven program delivery and strong community partnerships, RCHF is uniquely positioned to '
        'execute this plan and achieve lasting health improvements for thousands of underserved residents.'
    )

    doc.add_paragraph(
        'We respectfully request $245,000 in funding to bring this vision to reality. Thank you for '
        'your consideration of this application.'
    )

    # Add signature block
    doc.add_paragraph('')
    doc.add_paragraph('Respectfully submitted,')
    doc.add_paragraph('')
    doc.add_paragraph('Margaret Okonkwo, Executive Director')
    doc.add_paragraph('Riverside Community Health Foundation')
    doc.add_paragraph('March 1, 2026')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open with LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def _make_normal_run_elem(text):
    """Helper to create a plain w:r element."""
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


create_initial()
