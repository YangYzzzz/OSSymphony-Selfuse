"""
Initial Setup: Job posting document with misaligned columns (spaces, no tabs)
Task ID: writer_hr_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("TechVision Solutions - Open Positions", level=1)

    # Intro paragraph
    intro = doc.add_paragraph(
        "Below is a listing of current job openings across our offices. "
        "Please contact HR at hr@techvision.com for application details."
    )
    intro.paragraph_format.space_after = Pt(12)

    # Column header line - using spaces (deliberately misaligned)
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("Position          Salary Range          Location")
    header_run.bold = True
    header_run.font.size = Pt(11)
    header_run.font.name = "Calibri"
    header_para.paragraph_format.space_after = Pt(4)

    # Separator
    sep = doc.add_paragraph()
    sep_run = sep.add_run("=" * 65)
    sep_run.font.size = Pt(11)
    sep_run.font.name = "Calibri"
    sep.paragraph_format.space_after = Pt(4)

    # Job listings - spaces used for alignment (intentionally messy)
    jobs = [
        ("Senior Software Engineer", "$125,000 - $155,000", "San Francisco, CA"),
        ("Marketing Manager", "$85,000 - $105,000", "Austin, TX"),
        ("Data Analyst", "$72,000 - $92,000", "Chicago, IL"),
        ("UX Designer", "$90,000 - $115,000", "New York, NY"),
        ("DevOps Engineer", "$110,000 - $140,000", "Seattle, WA"),
        ("Product Manager", "$105,000 - $135,000", "San Francisco, CA"),
        ("HR Coordinator", "$55,000 - $68,000", "Denver, CO"),
        ("Financial Analyst", "$78,000 - $98,000", "Boston, MA"),
        ("QA Lead", "$95,000 - $120,000", "Portland, OR"),
        ("Technical Writer", "$65,000 - $82,000", "Austin, TX"),
        ("Sales Representative", "$60,000 - $75,000", "Miami, FL"),
        ("Cloud Architect", "$140,000 - $175,000", "Seattle, WA"),
    ]

    for position, salary, location in jobs:
        para = doc.add_paragraph()
        # Use spaces to simulate alignment attempt (creates misalignment)
        line = f"{position}     {salary}     {location}"
        run = para.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        para.paragraph_format.space_after = Pt(2)

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Last updated: March 2026. All positions are full-time with benefits."
    )
    footer_run = footer.runs[0]
    footer_run.italic = True
    footer_run.font.size = Pt(10)
    footer_run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
