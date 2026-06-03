"""
Reward Script: Set tab stops at 2 inches and 4 inches for job posting alignment
Task ID: writer_hr_024
Domain: libreoffice_writer

Scoring Rubric:
  Component 1 (0.35): Tab stops at 2 inches defined on data paragraphs
  Component 2 (0.35): Tab stops at 4 inches defined on data paragraphs
  Component 3 (0.30): Tab characters used in data paragraphs for column separation
"""

import os
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_024'

# Expected tab stop positions in EMU (English Metric Units)
# 2 inches = 2 * 914400 = 1828800 EMU
# 4 inches = 4 * 914400 = 3657600 EMU
TAB_2IN = 1828800
TAB_4IN = 3657600
# Allow 5% tolerance for tab positions
TOLERANCE = 0.05

# Data paragraphs are the header row (P2) and job listing rows (P4-P15)
# P0 = title, P1 = intro, P3 = separator, P16 = blank, P17 = footer
DATA_PARA_INDICES = [2, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15]


def get_effective_tab_stops(para):
    """Get non-default tab stops from a paragraph, filtering CLEAR and LEFT@0."""
    stops = []
    for ts in para.paragraph_format.tab_stops:
        if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
            continue
        if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
            continue
        stops.append(ts)
    return stops


def position_matches(actual_pos, expected_pos, tol=TOLERANCE):
    """Check if a tab stop position is within tolerance of expected."""
    return abs(actual_pos - expected_pos) / expected_pos <= tol


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has enough paragraphs
    if len(doc.paragraphs) < 16:
        print(f"FAIL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 16")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Tab stops at ~2 inches on data paragraphs (0.35 points)
    try:
        paras_with_2in = 0
        for idx in DATA_PARA_INDICES:
            para = doc.paragraphs[idx]
            stops = get_effective_tab_stops(para)
            has_2in = any(position_matches(ts.position, TAB_2IN) for ts in stops)
            if has_2in:
                paras_with_2in += 1

        ratio_2in = paras_with_2in / len(DATA_PARA_INDICES)
        if ratio_2in >= 0.8:
            print(f"PASS: Component 1 — Tab stop at 2in found on {paras_with_2in}/{len(DATA_PARA_INDICES)} data paragraphs (0.35 pts)")
            total_score += 0.35
        elif ratio_2in > 0:
            partial = round(0.35 * ratio_2in, 2)
            print(f"PARTIAL: Component 1 — Tab stop at 2in on {paras_with_2in}/{len(DATA_PARA_INDICES)} paragraphs ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No data paragraphs have tab stop at 2 inches")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Tab stops at ~4 inches on data paragraphs (0.35 points)
    try:
        paras_with_4in = 0
        for idx in DATA_PARA_INDICES:
            para = doc.paragraphs[idx]
            stops = get_effective_tab_stops(para)
            has_4in = any(position_matches(ts.position, TAB_4IN) for ts in stops)
            if has_4in:
                paras_with_4in += 1

        ratio_4in = paras_with_4in / len(DATA_PARA_INDICES)
        if ratio_4in >= 0.8:
            print(f"PASS: Component 2 — Tab stop at 4in found on {paras_with_4in}/{len(DATA_PARA_INDICES)} data paragraphs (0.35 pts)")
            total_score += 0.35
        elif ratio_4in > 0:
            partial = round(0.35 * ratio_4in, 2)
            print(f"PARTIAL: Component 2 — Tab stop at 4in on {paras_with_4in}/{len(DATA_PARA_INDICES)} paragraphs ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No data paragraphs have tab stop at 4 inches")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Tab characters used in data paragraphs (0.30 points)
    # In the golden state, spaces are replaced with tabs for alignment.
    # Each data paragraph should have at least 2 tab characters (separating 3 columns).
    try:
        paras_with_tabs = 0
        for idx in DATA_PARA_INDICES:
            para = doc.paragraphs[idx]
            tab_count = para.text.count('\t')
            if tab_count >= 2:
                paras_with_tabs += 1

        ratio_tabs = paras_with_tabs / len(DATA_PARA_INDICES)
        if ratio_tabs >= 0.8:
            print(f"PASS: Component 3 — Tab characters in {paras_with_tabs}/{len(DATA_PARA_INDICES)} data paragraphs (0.30 pts)")
            total_score += 0.30
        elif ratio_tabs > 0:
            partial = round(0.30 * ratio_tabs, 2)
            print(f"PARTIAL: Component 3 — Tab characters in {paras_with_tabs}/{len(DATA_PARA_INDICES)} paragraphs ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No data paragraphs use tab characters for alignment")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
