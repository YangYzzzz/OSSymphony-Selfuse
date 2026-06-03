"""
Initial Setup: Replace all soft hyphens in imported document with regular hyphens
Task ID: writer_frd_025
Domain: libreoffice_writer

Creates a realistic document simulating a PDF import that contains 30 soft hyphens
(\u00AD) scattered throughout compound words and hyphenated terms.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_025'
OUTPUT = f'{WORKDIR}/Imported_Text.docx'

SOFT_HYPHEN = '\u00AD'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # -- Title --
    title = doc.add_heading('Quarterly Performance Review: Northbridge Consulting Group', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- Subtitle --
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Fiscal Year 2024-2025 | Confidential')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')  # spacer

    # Soft hyphen helper: insert soft hyphen in a word
    # We need exactly 30 soft hyphens across the document
    # Using SH as marker, we'll place them in compound/hyphenated words

    # -- Section 1: Executive Summary (6 soft hyphens) --
    doc.add_heading('1. Executive Summary', level=2)

    p1 = doc.add_paragraph(
        f'This quarter{SOFT_HYPHEN}ly report provides a compre{SOFT_HYPHEN}hensive '
        f'overview of North{SOFT_HYPHEN}bridge Consulting Group\'s opera{SOFT_HYPHEN}tional '
        f'performance across all depart{SOFT_HYPHEN}ments. The manage{SOFT_HYPHEN}ment team has '
        f'identified several key areas of growth and strategic opportunities for the coming fiscal year.'
    )
    # Count: quarterly(1) comprehensive(2) Northbridge(3) operational(4) departments(5) management(6)

    # -- Section 2: Financial Overview (6 soft hyphens) --
    doc.add_heading('2. Financial Overview', level=2)

    p2 = doc.add_paragraph(
        f'Total revenue for the quarter reached $4.7 million, repre{SOFT_HYPHEN}senting '
        f'a 12% year{SOFT_HYPHEN}over{SOFT_HYPHEN}year increase. The profit{SOFT_HYPHEN}ability '
        f'margin improved to 18.3%, driven by cost{SOFT_HYPHEN}effective resource '
        f'allo{SOFT_HYPHEN}cation strategies implemented in Q2.'
    )
    # Count: representing(7) year-over-year(8,9) profitability(10) cost-effective(11) allocation(12)

    # -- Section 3: Department Performance (6 soft hyphens) --
    doc.add_heading('3. Department Performance', level=2)

    p3 = doc.add_paragraph(
        f'The techno{SOFT_HYPHEN}logy consulting division out{SOFT_HYPHEN}performed expectations '
        f'with a 23% increase in bill{SOFT_HYPHEN}able hours. Client satis{SOFT_HYPHEN}faction '
        f'scores remained consis{SOFT_HYPHEN}tently above 92% across all engage{SOFT_HYPHEN}ment types.'
    )
    # Count: technology(13) outperformed(14) billable(15) satisfaction(16) consistently(17) engagement(18)

    # -- Section 4: Human Resources (6 soft hyphens) --
    doc.add_heading('4. Human Resources', level=2)

    p4 = doc.add_paragraph(
        f'Employee reten{SOFT_HYPHEN}tion rates have stabi{SOFT_HYPHEN}lized at 94%, following '
        f'the imple{SOFT_HYPHEN}mentation of the new compen{SOFT_HYPHEN}sation framework. '
        f'The profes{SOFT_HYPHEN}sional develop{SOFT_HYPHEN}ment program enrolled 156 participants '
        f'this quarter, a record high for the organization.'
    )
    # Count: retention(19) stabilized(20) implementation(21) compensation(22) professional(23) development(24)

    # -- Section 5: Strategic Initiatives (6 soft hyphens) --
    doc.add_heading('5. Strategic Initiatives', level=2)

    p5 = doc.add_paragraph(
        f'The digital trans{SOFT_HYPHEN}formation initiative continues to progress ahead of '
        f'schedule. Infra{SOFT_HYPHEN}structure modern{SOFT_HYPHEN}ization efforts have reduced '
        f'system down{SOFT_HYPHEN}time by 67%. The inter{SOFT_HYPHEN}national expansion '
        f'strategy is under{SOFT_HYPHEN}going final review by the board of directors.'
    )
    # Count: transformation(25) Infrastructure(26) modernization(27) downtime(28) international(29) undergoing(30)

    # -- Closing --
    doc.add_paragraph('')
    closing = doc.add_paragraph(
        'This document was generated from an imported PDF source. '
        'Please review all formatting and special characters before distribution.'
    )
    closing_run = closing.runs[0]
    closing_run.font.italic = True
    closing_run.font.size = Pt(9)
    closing_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT)

    # Verify count
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    count = full_text.count(SOFT_HYPHEN)
    print(f'Initial file created: {OUTPUT}')
    print(f'Soft hyphen count: {count}')
    assert count == 30, f'Expected 30 soft hyphens, got {count}'

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
