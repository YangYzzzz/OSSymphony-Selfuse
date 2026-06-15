"""
Reward Script: Replace all soft hyphens with regular hyphens
Task ID: writer_frd_025
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Soft hyphens removed AND replaced with regular hyphens
  Component 2 (0.3): Correct total regular hyphen count (31 = 30 replacements + 1 original)
  Component 3 (0.2): Document structure preserved (paragraph count)

Note: LibreOffice strips soft hyphens on save without replacing them.
So we must verify replacement happened (regular hyphen count increased),
not just that soft hyphens are absent.
"""

import os


WORKDIR = '/home/user'
FILE_NAME = 'Imported_Text.docx'

SOFT_HYPHEN = chr(0x00AD)  # Unicode soft hyphen U+00AD

# Expected counts from initial document
EXPECTED_INITIAL_SOFT_HYPHENS = 30
EXPECTED_INITIAL_REGULAR_HYPHENS = 1
# After proper replacement: 30 soft hyphens -> 30 regular hyphens + 1 original = 31
EXPECTED_GOLDEN_REGULAR_HYPHENS = 31
EXPECTED_PARAGRAPH_COUNT = 15


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Count soft hyphens and regular hyphens across all paragraphs
    total_soft_hyphens = 0
    total_regular_hyphens = 0
    para_count = len(doc.paragraphs)

    for para in doc.paragraphs:
        text = para.text
        total_soft_hyphens += text.count(SOFT_HYPHEN)
        total_regular_hyphens += text.count('-')

    print(f"INFO: Found {total_soft_hyphens} soft hyphens, {total_regular_hyphens} regular hyphens, {para_count} paragraphs")

    # Component 1: Soft hyphens removed AND replaced with regular hyphens (0.5 points)
    # Both conditions must hold: no soft hyphens remain AND regular hyphen count increased
    # This prevents giving credit for simply opening/saving in LO (which strips soft hyphens
    # without adding regular hyphens)
    try:
        new_regular_hyphens = total_regular_hyphens - EXPECTED_INITIAL_REGULAR_HYPHENS
        if total_soft_hyphens == 0 and new_regular_hyphens >= EXPECTED_INITIAL_SOFT_HYPHENS:
            print(f"PASS: Component 1 -- Soft hyphens removed and replaced ({new_regular_hyphens} new regular hyphens) (0.5 pts)")
            total_score += 0.5
        elif total_soft_hyphens == 0 and new_regular_hyphens > 0:
            # Partial: some replacements done but not all
            partial = 0.5 * (new_regular_hyphens / EXPECTED_INITIAL_SOFT_HYPHENS)
            print(f"PARTIAL: Component 1 -- {new_regular_hyphens}/{EXPECTED_INITIAL_SOFT_HYPHENS} soft hyphens replaced ({partial:.2f} pts)")
            total_score += partial
        elif total_soft_hyphens < EXPECTED_INITIAL_SOFT_HYPHENS and new_regular_hyphens > 0:
            # Some soft hyphens removed and some replaced
            removed_ratio = (EXPECTED_INITIAL_SOFT_HYPHENS - total_soft_hyphens) / EXPECTED_INITIAL_SOFT_HYPHENS
            replaced_ratio = min(new_regular_hyphens / EXPECTED_INITIAL_SOFT_HYPHENS, 1.0)
            partial = 0.5 * min(removed_ratio, replaced_ratio)
            print(f"PARTIAL: Component 1 -- {total_soft_hyphens} soft hyphens remain, {new_regular_hyphens} new regular hyphens ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- soft_hyphens={total_soft_hyphens}, new_regular_hyphens={new_regular_hyphens}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Exact regular hyphen count matches expected (0.3 points)
    # Golden should have exactly 31 regular hyphens (30 replacements + 1 original)
    try:
        if total_regular_hyphens >= EXPECTED_GOLDEN_REGULAR_HYPHENS and total_soft_hyphens == 0:
            print(f"PASS: Component 2 -- Regular hyphen count correct ({total_regular_hyphens} >= {EXPECTED_GOLDEN_REGULAR_HYPHENS}) (0.3 pts)")
            total_score += 0.3
        elif total_regular_hyphens > EXPECTED_INITIAL_REGULAR_HYPHENS and total_soft_hyphens == 0:
            # Some replacements happened
            ratio = (total_regular_hyphens - EXPECTED_INITIAL_REGULAR_HYPHENS) / EXPECTED_INITIAL_SOFT_HYPHENS
            partial = 0.3 * min(ratio, 1.0)
            print(f"PARTIAL: Component 2 -- {total_regular_hyphens} regular hyphens (expected {EXPECTED_GOLDEN_REGULAR_HYPHENS}) ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- {total_regular_hyphens} regular hyphens (expected {EXPECTED_GOLDEN_REGULAR_HYPHENS}), soft_hyphens={total_soft_hyphens}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Document structure preserved AND replacement done (0.2 points)
    # Compound check: paragraph count is 15, text length reasonable, AND at least some
    # replacement occurred (new_regular_hyphens > 0). This anchors to the task change
    # so it fails on initial_env (where no replacement happened).
    try:
        new_regular_hyphens_c3 = total_regular_hyphens - EXPECTED_INITIAL_REGULAR_HYPHENS
        total_text = sum(len(p.text) for p in doc.paragraphs)
        if para_count == EXPECTED_PARAGRAPH_COUNT and total_text > 100 and new_regular_hyphens_c3 >= EXPECTED_INITIAL_SOFT_HYPHENS:
            print(f"PASS: Component 3 -- Document structure preserved with replacement ({para_count} paragraphs, {total_text} chars, {new_regular_hyphens_c3} replacements) (0.2 pts)")
            total_score += 0.2
        elif para_count != EXPECTED_PARAGRAPH_COUNT:
            print(f"FAIL: Component 3 -- Paragraph count changed: {para_count} (expected {EXPECTED_PARAGRAPH_COUNT})")
        elif total_text <= 100:
            print(f"FAIL: Component 3 -- Document appears corrupted (only {total_text} chars)")
        else:
            print(f"FAIL: Component 3 -- Structure OK but insufficient replacements ({new_regular_hyphens_c3} new regular hyphens)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point -- no persist hook; reading file as-is on disk
file_path = f'{WORKDIR}/{FILE_NAME}'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
