"""
Initial Setup: Copy the text from the second paragraph and paste it after the fourth paragraph
Task ID: writer_edit_062
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_062'
OUTPUT = f'{WORKDIR}/Desktop/training_manual.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document()

    # Remove default empty paragraph if present
    # Add the 5 paragraphs of the training manual
    paragraphs = [
        "Welcome to the training program.",
        "Safety is our top priority. All personnel must complete the safety orientation before accessing the facility.",
        "Training sessions are held every Monday and Wednesday.",
        "Please bring your employee ID badge to all sessions.",
        "Contact the training coordinator for schedule changes.",
    ]

    # Clear any default paragraphs and add our content
    for i, text in enumerate(paragraphs):
        if i == 0 and len(doc.paragraphs) > 0 and doc.paragraphs[0].text == '':
            # Reuse the first empty paragraph
            doc.paragraphs[0].add_run(text)
        else:
            doc.add_paragraph(text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
