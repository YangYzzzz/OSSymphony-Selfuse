"""
Reward script for writer_wf_081: IT Disaster Recovery Plan document verification.
Progressive scoring from 0.0 to 1.0.
"""
import os
import re

def compute_reward():
    score = 0.0
    total_weight = 0.0

    file_path = "/home/user/writer_wf_081.docx"

    # Gate: file must exist
    if not os.path.exists(file_path):
        print("File not found:", file_path)
        print("REWARD: 0.0")
        return

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print("Failed to open document:", e)
        print("REWARD: 0.0")
        return

    # Collect all paragraph info
    all_paras = []
    for p in doc.paragraphs:
        all_paras.append((p.style.name, p.text.strip()))

    heading1_texts = [t for s, t in all_paras if s == "Heading 1" and t]
    heading2_texts = [t for s, t in all_paras if s == "Heading 2" and t]
    all_text_lower = " ".join(t for _, t in all_paras).lower()

    # --- Component 1: Title (weight 0.10) ---
    w = 0.10
    total_weight += w
    try:
        title_paras = [t for s, t in all_paras if s == "Title" and t]
        title_text = " ".join(title_paras).lower()
        has_dr = "disaster recovery plan" in title_text
        has_nexora = "nexora technologies" in title_text
        if has_dr and has_nexora:
            score += w
            print("Title: PASS (both DR plan and Nexora Technologies found)")
        elif has_dr or has_nexora:
            score += w * 0.5
            print(f"Title: PARTIAL (DR plan: {has_dr}, Nexora: {has_nexora})")
        else:
            print("Title: FAIL")
    except Exception as e:
        print(f"Title: ERROR - {e}")

    # --- Component 2: Document Control Info (weight 0.10) ---
    w = 0.10
    total_weight += w
    try:
        doc_control_score = 0.0
        # Check for document control table or text
        # The golden has a "Document Control" heading and a table with Version, Date, Author, Classification
        has_version = False
        has_date = False
        has_confidential = False

        # Check in tables
        for table in doc.tables:
            for row in table.rows:
                cells_text = [c.text.strip().lower() for c in row.cells]
                row_text = " ".join(cells_text)
                if "version" in row_text:
                    has_version = True
                if "date" in row_text:
                    has_date = True
                if "confidential" in row_text:
                    has_confidential = True

        # Also check in paragraph text
        if not has_version:
            has_version = "version" in all_text_lower
        if not has_date:
            has_date = bool(re.search(r'\d{4}[-/]\d{2}[-/]\d{2}', " ".join(t for _, t in all_paras)))
        if not has_confidential:
            has_confidential = "confidential" in all_text_lower

        items_found = sum([has_version, has_date, has_confidential])
        doc_control_score = items_found / 3.0
        score += w * doc_control_score
        print(f"Document Control: {items_found}/3 (version={has_version}, date={has_date}, confidential={has_confidential})")
    except Exception as e:
        print(f"Document Control: ERROR - {e}")

    # --- Component 3: Table of Contents (weight 0.05) ---
    w = 0.05
    total_weight += w
    try:
        has_toc = False
        for s, t in all_paras:
            tl = t.lower()
            if "table of contents" in tl or (s == "Heading 1" and "contents" in tl):
                has_toc = True
                break
        # Also check if there's a TOC-like listing (numbered section references)
        if not has_toc:
            toc_pattern_count = 0
            for _, t in all_paras:
                if re.match(r'^\d+\.\s+\w', t.strip()):
                    toc_pattern_count += 1
            if toc_pattern_count >= 4:
                has_toc = True

        if has_toc:
            score += w
            print("TOC: PASS")
        else:
            print("TOC: FAIL")
    except Exception as e:
        print(f"TOC: ERROR - {e}")

    # --- Component 4: 6 Heading 1 sections (weight 0.15) ---
    w = 0.15
    total_weight += w
    try:
        expected_h1 = [
            "plan overview",
            "recovery team",
            "risk scenarios",
            "recovery procedures",
            "testing schedule",
            "plan maintenance"
        ]
        found_h1 = 0
        h1_lower = [h.lower() for h in heading1_texts]
        for exp in expected_h1:
            for h in h1_lower:
                if exp in h:
                    found_h1 += 1
                    break
        h1_ratio = min(found_h1 / 6.0, 1.0)
        score += w * h1_ratio
        print(f"Heading 1 sections: {found_h1}/6 found (score ratio: {h1_ratio:.2f})")
    except Exception as e:
        print(f"Heading 1 sections: ERROR - {e}")

    # --- Component 5: Recovery Team table (weight 0.15) ---
    w = 0.15
    total_weight += w
    try:
        rt_score = 0.0
        best_rt_score = 0.0
        for table in doc.tables:
            if len(table.rows) < 2 or len(table.columns) < 3:
                continue
            headers = [c.text.strip().lower() for c in table.rows[0].cells]
            header_text = " ".join(headers)
            # Check if this looks like the recovery team table
            has_role = any("role" in h for h in headers)
            has_name = any("name" in h for h in headers)
            has_contact = any("contact" in h for h in headers)
            has_backup = any("backup" in h for h in headers)

            if has_role and has_name:
                col_score = sum([has_role, has_name, has_contact, has_backup]) / 4.0
                data_rows = len(table.rows) - 1  # exclude header
                row_score = min(data_rows / 5.0, 1.0)
                table_score = 0.5 * col_score + 0.5 * row_score
                best_rt_score = max(best_rt_score, table_score)

        score += w * best_rt_score
        print(f"Recovery Team table: score={best_rt_score:.2f}")
    except Exception as e:
        print(f"Recovery Team table: ERROR - {e}")

    # --- Component 6: Risk Scenarios table (weight 0.15) ---
    w = 0.15
    total_weight += w
    try:
        best_rs_score = 0.0
        for table in doc.tables:
            if len(table.rows) < 2 or len(table.columns) < 3:
                continue
            headers = [c.text.strip().lower() for c in table.rows[0].cells]
            has_scenario = any("scenario" in h for h in headers)
            has_rto = any("rto" in h for h in headers)
            has_rpo = any("rpo" in h for h in headers)
            has_priority = any("priority" in h for h in headers)

            if has_scenario and (has_rto or has_rpo):
                col_score = sum([has_scenario, has_rto, has_rpo, has_priority]) / 4.0
                data_rows = len(table.rows) - 1
                row_score = min(data_rows / 4.0, 1.0)
                table_score = 0.5 * col_score + 0.5 * row_score
                best_rs_score = max(best_rs_score, table_score)

        score += w * best_rs_score
        print(f"Risk Scenarios table: score={best_rs_score:.2f}")
    except Exception as e:
        print(f"Risk Scenarios table: ERROR - {e}")

    # --- Component 7: 4 Heading 2 sub-procedures (weight 0.10) ---
    w = 0.10
    total_weight += w
    try:
        # Check for Heading 2 entries under Recovery Procedures
        # Look for heading 2s that appear after "Recovery Procedures" heading 1
        in_recovery = False
        h2_count = 0
        for s, t in all_paras:
            if s == "Heading 1" and "recovery procedures" in t.lower():
                in_recovery = True
                continue
            if s == "Heading 1" and in_recovery:
                break  # next H1, stop counting
            if in_recovery and s == "Heading 2" and t:
                h2_count += 1

        # Fallback: just count all Heading 2s if none found in context
        if h2_count == 0:
            h2_count = len(heading2_texts)

        h2_ratio = min(h2_count / 4.0, 1.0)
        score += w * h2_ratio
        print(f"Heading 2 sub-procedures: {h2_count}/4 found (score ratio: {h2_ratio:.2f})")
    except Exception as e:
        print(f"Heading 2 sub-procedures: ERROR - {e}")

    # --- Component 8: Testing Schedule table (weight 0.10) ---
    w = 0.10
    total_weight += w
    try:
        best_ts_score = 0.0
        for table in doc.tables:
            if len(table.rows) < 2 or len(table.columns) < 2:
                continue
            headers = [c.text.strip().lower() for c in table.rows[0].cells]
            has_test_type = any("test" in h and "type" in h for h in headers)
            has_frequency = any("frequency" in h or "freq" in h for h in headers)
            has_last_tested = any("last" in h and "test" in h for h in headers)

            if has_test_type or (has_frequency and has_last_tested):
                col_score = sum([has_test_type, has_frequency, has_last_tested]) / 3.0
                data_rows = len(table.rows) - 1
                row_score = min(data_rows / 3.0, 1.0)
                table_score = 0.5 * col_score + 0.5 * row_score
                best_ts_score = max(best_ts_score, table_score)

        score += w * best_ts_score
        print(f"Testing Schedule table: score={best_ts_score:.2f}")
    except Exception as e:
        print(f"Testing Schedule table: ERROR - {e}")

    # --- Component 9: Plan Maintenance content (weight 0.10) ---
    w = 0.10
    total_weight += w
    try:
        in_maintenance = False
        maintenance_content = []
        for s, t in all_paras:
            if s == "Heading 1" and "plan maintenance" in t.lower():
                in_maintenance = True
                continue
            if s == "Heading 1" and in_maintenance:
                break
            if in_maintenance and t:
                maintenance_content.append(t)

        if len(maintenance_content) >= 1 and len(" ".join(maintenance_content)) > 20:
            score += w
            print(f"Plan Maintenance: PASS ({len(maintenance_content)} paragraphs)")
        elif len(maintenance_content) >= 1:
            score += w * 0.5
            print(f"Plan Maintenance: PARTIAL ({len(maintenance_content)} paragraphs, short)")
        else:
            print("Plan Maintenance: FAIL (no content)")
    except Exception as e:
        print(f"Plan Maintenance: ERROR - {e}")

    # Final score
    final_score = round(min(max(score, 0.0), 1.0), 1)
    print(f"\nTotal weight checked: {total_weight:.2f}")
    print(f"Raw score: {score:.4f}")
    print(f"REWARD: {final_score}")

if __name__ == "__main__":
    compute_reward()
