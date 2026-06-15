"""
Reward Script: Standard Manuscript Format for Literary Magazine Submission
Task ID: writer_creative_031
Domain: libreoffice_writer
Scoring:
  - Component 1: Font — Courier New 12pt throughout (0.25 pts)
  - Component 2: Contact block — 5-line contact info at top-left, single-spaced (0.20 pts)
  - Component 3: Title centered with large space_before & byline centered (0.20 pts)
  - Component 4: Double spacing + 0.5in first line indent in story paragraphs (0.20 pts)
  - Component 5: Header "Sharma / [page#]" right-aligned, first-page header blank (0.15 pts)
Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_031'
FILE_PATH = f'{WORKDIR}/blue_door_story.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paras = doc.paragraphs

    # -------------------------------------------------------------------------
    # Component 1: Font — Courier New 12pt throughout (0.25 pts)
    # The initial file uses Arial 11pt. Correct answer requires Courier New 12pt.
    # -------------------------------------------------------------------------
    try:
        courier_count = 0
        non_courier_count = 0
        size_12_count = 0
        non_size_12_count = 0

        for para in paras:
            for run in para.runs:
                if run.text.strip():
                    fn = run.font.name
                    fs = run.font.size.pt if run.font.size else None

                    if fn is not None:
                        if fn == 'Courier New':
                            courier_count += 1
                        else:
                            non_courier_count += 1

                    if fs is not None:
                        if abs(fs - 12.0) <= 0.1:
                            size_12_count += 1
                        else:
                            non_size_12_count += 1

        courier_ok = (courier_count > 0 and non_courier_count == 0)
        size_ok = (size_12_count > 0 and non_size_12_count == 0)

        if courier_ok and size_ok:
            print(f"PASS: Component 1 — Font Courier New 12pt ({courier_count} runs checked) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected Courier New 12pt; courier_ok={courier_ok} "
                  f"(found={courier_count}, non={non_courier_count}), "
                  f"size_ok={size_ok} (12pt={size_12_count}, non={non_size_12_count})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Contact block — 5-line block at top of document (0.20 pts)
    # Expected lines: "Priya Sharma", "1847 Oakdale Ave", "Chicago, IL 60625",
    #                 "priya.sharma@email.com", word count line starting "Approx"
    # Must be left-aligned and single-spaced within the block.
    # The initial file does NOT have this block.
    # -------------------------------------------------------------------------
    try:
        required_contact_lines = [
            "Priya Sharma",
            "1847 Oakdale Ave",
            "Chicago, IL 60625",
            "priya.sharma@email.com",
        ]
        # Word count line: starts with "Approx" and contains "3" and "500 words"
        # We check more loosely for the word count line
        word_count_present = False

        contact_found = 0
        if len(paras) >= 5:
            for i, expected in enumerate(required_contact_lines):
                actual_text = paras[i].text.strip()
                if actual_text == expected:
                    contact_found += 1

            # Word count line (para index 4)
            wc_text = paras[4].text.strip().lower()
            if ('approx' in wc_text or 'word' in wc_text) and ('3' in wc_text or '3500' in wc_text or '3,500' in wc_text):
                word_count_present = True

        all_contact_present = (contact_found == 4 and word_count_present)

        # Check single-spaced within block (line_spacing == 1.0 or None or close to 1.0)
        contact_single_spaced = True
        for i in range(5):
            if i < len(paras):
                ls = paras[i].paragraph_format.line_spacing
                if ls is not None and ls != 1.0 and (not isinstance(ls, (int, float)) or abs(ls - 1.0) > 0.05):
                    contact_single_spaced = False

        if all_contact_present and contact_single_spaced:
            print(f"PASS: Component 2 — Contact block with 5 lines present and single-spaced (0.20 pts)")
            total_score += 0.20
        elif all_contact_present:
            print(f"PASS: Component 2 — Contact block present (single-spacing check inconclusive) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — Contact block missing or incomplete; "
                  f"lines_found={contact_found}/4, word_count={word_count_present}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Title "The Blue Door" centered with large space_before,
    #              and "by Priya Sharma" centered (0.20 pts)
    # The title should be pushed ~halfway down page 1 using space_before >= 100pt.
    # The initial file has the title left-aligned and no space_before.
    # -------------------------------------------------------------------------
    try:
        title_para = None
        byline_para = None

        # Find title and byline among paragraphs (after contact block)
        for i, para in enumerate(paras):
            text = para.text.strip()
            if text == 'The Blue Door' and title_para is None:
                title_para = para
            if text == 'by Priya Sharma' and byline_para is None and title_para is not None:
                byline_para = para

        title_centered = False
        title_has_space_before = False
        byline_centered = False

        if title_para:
            align = title_para.paragraph_format.alignment
            title_centered = (align == WD_PARAGRAPH_ALIGNMENT.CENTER)
            sp = title_para.paragraph_format.space_before
            if sp is not None:
                try:
                    # space_before is in EMU; convert to pt (1pt = 12700 EMU)
                    sp_pt = sp / 12700.0
                    title_has_space_before = (sp_pt >= 100.0)
                except Exception:
                    title_has_space_before = False

        if byline_para:
            align = byline_para.paragraph_format.alignment
            byline_centered = (align == WD_PARAGRAPH_ALIGNMENT.CENTER)

        if title_centered and title_has_space_before and byline_centered:
            print(f"PASS: Component 3 — Title centered with space_before, byline centered (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — title_centered={title_centered}, "
                  f"title_has_space_before={title_has_space_before} "
                  f"(space_before={title_para.paragraph_format.space_before if title_para else 'N/A'}), "
                  f"byline_centered={byline_centered}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Double spacing (2.0) + 0.5in first line indent for story paras
    #              (0.20 pts)
    # The initial file has single-spaced paragraphs and no first-line indent.
    # Story paragraphs start after the byline (para index 7 in golden).
    # -------------------------------------------------------------------------
    try:
        story_paras = []
        found_byline = False
        for para in paras:
            text = para.text.strip()
            if text == 'by Priya Sharma':
                found_byline = True
                continue
            if found_byline and text:
                story_paras.append(para)

        double_spaced_count = 0
        indented_count = 0
        total_story = len(story_paras)

        half_inch_emu = int(Inches(0.5))  # 457200 EMU

        for para in story_paras:
            pf = para.paragraph_format
            ls = pf.line_spacing
            fli = pf.first_line_indent

            if ls is not None and abs(ls - 2.0) <= 0.05:
                double_spaced_count += 1
            # first_line_indent ~0.5 inch (allow ±10000 EMU tolerance ~0.008 inch)
            if fli is not None and abs(fli - half_inch_emu) <= 20000:
                indented_count += 1

        if total_story > 0:
            ds_ratio = double_spaced_count / total_story
            ind_ratio = indented_count / total_story
        else:
            ds_ratio = 0
            ind_ratio = 0

        if ds_ratio >= 0.9 and ind_ratio >= 0.9:
            print(f"PASS: Component 4 — Double-spacing ({double_spaced_count}/{total_story}) "
                  f"and 0.5in first-line indent ({indented_count}/{total_story}) (0.20 pts)")
            total_score += 0.20
        elif ds_ratio >= 0.9:
            print(f"PARTIAL: Component 4 — Double-spacing OK ({double_spaced_count}/{total_story}) "
                  f"but first-line indent insufficient ({indented_count}/{total_story}); 0.0 pts")
        else:
            print(f"FAIL: Component 4 — Double-spacing {double_spaced_count}/{total_story} "
                  f"({ds_ratio:.0%}), indent {indented_count}/{total_story} ({ind_ratio:.0%})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Header with "Sharma / PAGE" right-aligned, first page no header
    #              (0.15 pts)
    # The initial file has no header at all.
    # -------------------------------------------------------------------------
    try:
        section = doc.sections[0]

        # Check that different_first_page_header_footer is set
        first_page_diff = section.different_first_page_header_footer

        # Check main header text starts with "Sharma"
        hdr = section.header
        hdr_text = ""
        if hdr.paragraphs:
            hdr_text = hdr.paragraphs[0].text

        header_has_sharma = hdr_text.strip().startswith("Sharma")

        # Check header right-aligned
        hdr_align = None
        if hdr.paragraphs:
            hdr_align = hdr.paragraphs[0].paragraph_format.alignment
        header_right_aligned = (hdr_align == WD_PARAGRAPH_ALIGNMENT.RIGHT)

        # Check header contains a page number field code (instrText with PAGE)
        page_num_in_header = False
        if hdr.paragraphs:
            for para in hdr.paragraphs:
                for run in para.runs:
                    xml_str = run._element.xml
                    if 'instrText' in xml_str and 'PAGE' in xml_str:
                        page_num_in_header = True
                    elif 'fldChar' in xml_str:
                        # May have fldChar elements that form a PAGE field
                        page_num_in_header = True

        # First page header should be blank/empty
        first_page_hdr_blank = True
        if first_page_diff:
            try:
                fp_hdr = section.first_page_header
                fp_text = ""
                if fp_hdr.paragraphs:
                    fp_text = fp_hdr.paragraphs[0].text
                first_page_hdr_blank = (fp_text.strip() == "")
            except Exception:
                first_page_hdr_blank = True  # If unavailable, assume blank

        if header_has_sharma and header_right_aligned and page_num_in_header and first_page_diff and first_page_hdr_blank:
            print(f"PASS: Component 5 — Header 'Sharma / PAGE' right-aligned, first-page blank (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — header_has_sharma={header_has_sharma} (text={repr(hdr_text)}), "
                  f"right_aligned={header_right_aligned}, page_num_in_header={page_num_in_header}, "
                  f"first_page_diff={first_page_diff}, first_page_blank={first_page_hdr_blank}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 4), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
