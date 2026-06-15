"""
Initial Setup: Short story manuscript before formatting (blue_door_story.docx)
Task ID: writer_creative_031
Domain: libreoffice_writer

Creates a short story file in ~/Desktop/ with:
- Title and byline
- 12 paragraphs of story content (~3500 words)
- 11pt Arial, single-spaced, no headers, no contact block, no indents
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'blue_door_story'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default font to Arial 11pt (single-spaced, no margins changes)
    section = doc.sections[0]
    # Leave default margins (not 1-inch — the task asks the agent to change them)
    # Default python-docx margins are 1.25 inches; leave as is to be different from task goal

    # Set document default font style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(8)

    # Title line 1
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_para.paragraph_format.line_spacing = 1.0
    title_para.paragraph_format.space_after = Pt(2)
    run = title_para.add_run('The Blue Door')
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    # Byline line 2
    byline_para = doc.add_paragraph()
    byline_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    byline_para.paragraph_format.line_spacing = 1.0
    byline_para.paragraph_format.space_after = Pt(8)
    run = byline_para.add_run('by Priya Sharma')
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    # 12 paragraphs of story text (~3500 words total)
    story_paragraphs = [
        # Paragraph 1 (~290 words)
        "The door at the end of Maplewood Lane had always been blue. Not the cheerful turquoise of beach houses, "
        "nor the dusty navy of old government buildings, but a deep, resonant blue—the color of the sky just "
        "before dawn surrenders to daylight. Nadia Petrov had walked past it every morning for eleven years, "
        "ever since her family moved into the rental two blocks away. She knew every chip in its paint, every "
        "scrape along the lower corner where someone had once dragged a bicycle carelessly. She knew that the "
        "brass knocker, shaped like a fox with a curling tail, always caught the light at seven forty-five in "
        "the morning in October. She knew that lavender grew in the narrow strip of garden beside the steps, "
        "and that in June the whole street smelled faintly of it. What she did not know—had never considered, "
        "despite eleven years of evidence—was that the door would one day open for her.",

        # Paragraph 2 (~300 words)
        "It happened on a Thursday, which seemed wrong. Thursdays were ordinary. Thursdays smelled like the "
        "corner bakery's second batch of bread, the one that came out slightly too dark and filled the morning "
        "air with a toasty, almost burnt sweetness. Thursdays meant her colleague Marcus would bring in his "
        "leftover curry for lunch and the entire archive floor would smell of cumin until closing. Thursdays "
        "were not for doors opening that had never opened before. And yet, as Nadia passed the blue door at "
        "precisely seven forty-three—running two minutes behind because she had stopped to rescue a sparrow "
        "that had stunned itself against her window—the door swung inward with a low, melodic creak. An old "
        "woman stood in the frame, her white hair pinned high, wearing a housecoat the precise grey of winter "
        "clouds. She looked directly at Nadia. 'You walk past every day,' the woman said. It was not a "
        "question. Nadia stopped on the pavement, her bag half-slipping off one shoulder. 'I do,' she agreed, "
        "because what else could she say. 'Come in, then,' the woman said. 'I have been waiting long enough.'",

        # Paragraph 3 (~290 words)
        "The inside of the house was nothing like Nadia had imagined. She realized, stepping over the "
        "threshold, that she had imagined it—had constructed in her mind over years a version of what such "
        "a house must contain. She had pictured dim rooms crowded with old furniture, perhaps a grandfather "
        "clock ticking in a hallway, perhaps cats. There were no cats. There was no grandfather clock. "
        "Instead the hallway was bright and spare, painted the same startling blue as the door, hung with "
        "a dozen small paintings, each no larger than a paperback book. They were all of the same subject: "
        "a window, seen from the inside, with different weathers beyond it. Rain. Snow. A cloudless summer "
        "day. Fog so thick the glass seemed painted white. Each painting had a tiny handwritten label: a "
        "single word, always in a language Nadia did not recognize. The old woman moved down the hall "
        "without explaining, and Nadia followed because she had already committed to this, had already "
        "crossed the threshold, and turning back now seemed both rude and somehow impossible.",

        # Paragraph 4 (~295 words)
        "The kitchen was at the end of the hall, and it was large and warm and smelled of cardamom and "
        "something floral Nadia couldn't name. A kettle was already singing on the stove. The table was "
        "round and set for two, which should have been alarming—two cups, two small plates, a plate of "
        "biscuits dusted with icing sugar—but instead felt, peculiarly, like relief. 'Sit,' the woman "
        "said, and Nadia sat. 'My name is Elena,' the woman said, pouring water into a teapot without "
        "looking. 'Elena Vaszary. I came here in 1971 from a town you have never heard of. I have lived "
        "in this house for fifty-two years and you are the first person who has ever come inside it.' "
        "She set the teapot on the table and sat down across from Nadia and looked at her with eyes "
        "that were very dark and very calm, the eyes of someone who has long since stopped being "
        "surprised by the world but has not stopped being interested in it.",

        # Paragraph 5 (~295 words)
        "Nadia asked how that was possible. Fifty-two years in one house with no visitors. Elena "
        "smiled—not a sad smile, more the smile of someone correcting a small misunderstanding. "
        "'I had visitors,' she said. 'I did not want them. There is a difference.' She poured the "
        "tea, a pale amber color, and pushed the cup toward Nadia. 'You are different. I have "
        "watched you watch the door. Not the way people look at something they want. The way "
        "people look at something they are trying to understand.' Nadia wrapped her hands around "
        "the cup and thought about this. It was true. She had never wanted to go inside. She had "
        "never fantasized about what the house contained or who lived there. She had simply "
        "noticed, every morning, that the door was blue and that this blue seemed to mean "
        "something, though she could not have said what. 'I am a librarian,' Nadia said finally, "
        "because it seemed relevant in a way she couldn't articulate. Elena nodded as though "
        "this confirmed something she had already suspected.",

        # Paragraph 6 (~290 words)
        "They drank the tea. Elena talked about the town she had come from—not its name, but "
        "its character. The river that divided it. The market that ran every Friday and smelled "
        "of beets and woodsmoke. The way winter arrived there not gradually but all at once, "
        "overnight, so that you could wake up in autumn and step outside into full winter. Nadia "
        "talked about the archive where she worked, the collection of local newspapers going "
        "back to 1887, the smell of old newsprint that she had stopped noticing years ago but "
        "that visitors always mentioned with either pleasure or alarm. They talked about the "
        "sparrow—Elena wanted to know if it had been injured or just dazed. Just dazed, Nadia "
        "said. Elena nodded. 'The dazed ones usually recover,' she said. 'They fly away "
        "confused for a day or two and then they are fine. It is the same with people.' "
        "Nadia ate a biscuit and didn't ask what she meant.",

        # Paragraph 7 (~290 words)
        "After an hour Nadia realized she was late for work and said so, and Elena rose to "
        "show her out without any protest or urging her to stay. At the door Elena said, "
        "'Come again if you like. I am always here in the morning. The door will be open.' "
        "Nadia walked the remaining two blocks to the archive in a state of mild disorientation. "
        "Marcus was already in his usual chair when she arrived, his coat still on, eating a "
        "pastry and reading something on his phone. 'You're late,' he said without looking up. "
        "'I know,' Nadia said. She sat down, turned on her monitor, and tried to think about "
        "the digitization backlog she was supposed to be processing. She thought instead about "
        "the small paintings in the hallway and the word on each one in the unknown language "
        "and what the word for fog might mean in a place where fog could arrive without warning.",

        # Paragraph 8 (~290 words)
        "She went back the following Thursday. Elena opened the door before she knocked, "
        "which Nadia chose not to find strange. The tea was already made. This time there "
        "were small sandwiches as well. Elena had a box of photographs on the table—black "
        "and white, many of them, some very old. 'I have been organizing,' she said, which "
        "was how they spent the next hour, Elena placing photographs into groupings she "
        "seemed to have already decided on and explaining each one: who these people were, "
        "where this square was, what year this house had been built. Nadia listened and asked "
        "questions when she had them, which was often. She was good at listening. It was "
        "most of what her job required—listening to collections tell their story, attending "
        "to what the documents said and what they left out. By the time she left she felt "
        "she knew the river, the market, the town whose name Elena still hadn't given her.",

        # Paragraph 9 (~290 words)
        "On the third visit Elena showed her the paintings. They went into a small back room "
        "Nadia hadn't seen before, which was where the rest of the windows lived—not twelve "
        "but forty-seven, Elena said, painted over thirty years, each one a window in a "
        "different place. Elena herself had painted them. She had studied painting as a young "
        "woman, before the river and the market and the winter arriving overnight. 'I only "
        "paint windows,' she said, standing in the center of the room with her hands folded. "
        "'Other people paint what they see through windows. I paint the window itself. The "
        "frame. The glass. The light the glass holds.' Nadia looked at them a long time. "
        "Each one was slightly different in some quality she couldn't name—not the weather "
        "beyond, which obviously varied, but something in the glass itself, some quality of "
        "attention. As though the glass were doing something. 'What are the words?' she "
        "asked. Elena looked at her steadily. 'They are not words for things,' she said. "
        "'They are words for kinds of looking.'",

        # Paragraph 10 (~295 words)
        "Winter came and they kept meeting on Thursdays. Sometimes Elena told stories and "
        "sometimes Nadia did. She told Elena about the archive's strangest holdings—the "
        "collection of letters from a man who had written to the town council every week "
        "for seventeen years complaining about a specific street corner, the letters "
        "becoming increasingly elaborate and eventually almost beautiful, the final one "
        "a kind of prose poem about pavement. Elena laughed at this, a full and genuine "
        "laugh that seemed to surprise her. She told Nadia about the year the river had "
        "flooded, which had happened twice in her memory: the first time devastating, the "
        "second time, which she had been young for, like a festival, the whole town wading "
        "through knee-deep water in borrowed rubber boots, lending each other things, "
        "making meals on second floors, temporarily becoming the kind of community that "
        "usually only exists in memory or fiction.",

        # Paragraph 11 (~295 words)
        "In February Elena gave her a painting. It was small like the others, a window in "
        "a stone wall with a cold white sky beyond it and frost patterns in the corner of "
        "the glass, and the single word in the unknown language on the lower right. "
        "'What does this one say?' Nadia asked. Elena considered. 'Approximately,' she "
        "said, 'it means: the quality of light before a thing you have been waiting for "
        "arrives.' She paused. 'It does not translate perfectly.' Nadia held the painting "
        "carefully and looked at the frost and the white sky and thought about eleven "
        "years of walking past a blue door. She thought about the sparrow stunned against "
        "her window, dazed for a day and then fine, flying away confused and then not "
        "confused. She thought about the word for fog and all the other words on all "
        "the other paintings. She thought: I have been waiting a long time and I did not "
        "know it. She thought: this is what the blue has always been trying to say.",

        # Paragraph 12 (~290 words)
        "She hung the painting in her apartment, above the desk where she read in the "
        "evenings. Visitors asked about it—the unusual quality of the light, the frost "
        "patterns, the word they couldn't read. She told them it was a gift from a friend "
        "who painted windows and had a blue door on Maplewood Lane. She did not try to "
        "explain the word because Elena had told her it didn't translate perfectly, and "
        "Nadia, who had spent her career attending to what things said and what they left "
        "out, understood that some meanings live in the original and lose something "
        "irretrievable in translation. The frost stayed perfect. The white sky stayed "
        "waiting. On Thursday mornings when she passed the blue door it was always open "
        "a little now, just a crack, the way doors are left when someone inside wants you "
        "to know you are expected. She always went in.",
    ]

    for para_text in story_paragraphs:
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.first_line_indent = None  # No indent in initial state
        run = para.add_run(para_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
