"""
Initial Setup: Set paragraph spacing for an Employee Handbook document.
Task ID: wrpara_021
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_021'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title: Employee Handbook (Heading 1 style) ---
    title = doc.add_heading('Employee Handbook', level=1)
    # Ensure all spacing is 0
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)

    # === Section 1: Attendance ===
    h1 = doc.add_heading('Attendance', level=2)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)

    p1 = doc.add_paragraph(
        'All employees are expected to arrive at their designated workstation by 9:00 AM '
        'and remain until 5:30 PM, Monday through Friday. Tardiness of more than 15 minutes '
        'will be recorded and reviewed by the HR department on a quarterly basis.'
    )
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)

    p2 = doc.add_paragraph(
        'Employees who anticipate a late arrival or early departure must notify their direct '
        'supervisor at least two hours in advance. Unexcused absences exceeding three days in '
        'a rolling 90-day period may result in a formal performance review.'
    )
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    # === Section 2: Dress Code ===
    h2 = doc.add_heading('Dress Code', level=2)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)

    p3 = doc.add_paragraph(
        'Business casual attire is required in all client-facing areas of the office. This '
        'includes collared shirts, slacks or skirts of appropriate length, and closed-toe shoes. '
        'Denim jeans are permitted on Fridays only, provided they are in good condition.'
    )
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)

    p4 = doc.add_paragraph(
        'Employees working in the warehouse or laboratory facilities must wear company-issued '
        'safety gear at all times, including steel-toe boots, high-visibility vests, and '
        'protective eyewear. Failure to comply may result in restricted facility access.'
    )
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(0)

    # === Section 3: Remote Work ===
    h3 = doc.add_heading('Remote Work', level=2)
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)

    p5 = doc.add_paragraph(
        'Eligible employees may work remotely up to two days per week with prior manager approval. '
        'Remote work requests must be submitted through the HR portal by the 15th of the preceding '
        'month. Approval is contingent upon satisfactory performance reviews and role suitability.'
    )
    p5.paragraph_format.space_before = Pt(0)
    p5.paragraph_format.space_after = Pt(0)

    p6 = doc.add_paragraph(
        'During remote work days, employees must remain available via company communication channels '
        'between 9:00 AM and 5:00 PM. A stable internet connection and a dedicated workspace are '
        'required. The company reserves the right to revoke remote work privileges with 14 days notice.'
    )
    p6.paragraph_format.space_before = Pt(0)
    p6.paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
