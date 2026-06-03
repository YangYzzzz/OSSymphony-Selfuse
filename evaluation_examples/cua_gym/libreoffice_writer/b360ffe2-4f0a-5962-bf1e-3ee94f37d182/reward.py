"""
Reward Script: Convert marketing budget narrative to structured document
Task ID: writer_mktg_027
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Budget table with 3 columns, 7 rows (header + 5 data + total)
  Component 2 (0.20): Table data rows contain correct budget items with currency amounts and % of total
  Component 3 (0.15): Total row bold with $500,000
  Component 4 (0.20): 4 numbered list items (List Number style) with correct action items
  Component 5 (0.10): "Budget Summary" and "Key Action Items" headings are bold and 14pt
Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_027'
FILE_PATH = f'{WORKDIR}/budget_justification_2026.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must exist and be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Budget table exists with correct dimensions (0.35 points)
    # The task asks for a table with 3 columns (Line Item, Budget Amount, % of Total)
    # and 7 rows (1 header + 5 data rows + 1 total row)
    try:
        tables = doc.tables
        budget_table = None

        # Find the budget table by looking for one with expected header
        for table in tables:
            if len(table.rows) > 0 and len(table.columns) == 3:
                header_texts = [c.text.strip().lower() for c in table.rows[0].cells]
                if any('line item' in h or 'budget' in h or '% of total' in h for h in header_texts):
                    budget_table = table
                    break

        if budget_table is None:
            print(f"FAIL: Component 1 — No budget table with 3 columns found. Tables found: {len(tables)}")
        elif len(budget_table.rows) != 7:
            print(f"FAIL: Component 1 — Table found but has {len(budget_table.rows)} rows, expected 7 (1 header + 5 data + 1 total)")
        elif len(budget_table.columns) != 3:
            print(f"FAIL: Component 1 — Table has {len(budget_table.columns)} columns, expected 3")
        else:
            # Check correct column headers
            header_row = budget_table.rows[0]
            col0 = header_row.cells[0].text.strip()
            col1 = header_row.cells[1].text.strip()
            col2 = header_row.cells[2].text.strip()
            if ('line item' in col0.lower() and 'budget' in col1.lower() and '%' in col2.lower()):
                print(f"PASS: Component 1 — Budget table found with correct dimensions 7x3 and headers: {col0!r}, {col1!r}, {col2!r} (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 1 — Table headers incorrect: {col0!r}, {col1!r}, {col2!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table data rows have correct budget items and % of total values (0.20 points)
    # Expected: 5 data rows with specific items, amounts, and percentages
    try:
        tables = doc.tables
        budget_table = None
        for table in tables:
            if len(table.rows) >= 7 and len(table.columns) == 3:
                header_texts = [c.text.strip().lower() for c in table.rows[0].cells]
                if any('line item' in h or 'budget' in h or '% of total' in h for h in header_texts):
                    budget_table = table
                    break

        if budget_table is None:
            print(f"FAIL: Component 2 — No budget table found for data validation")
        else:
            # Expected data rows: item name, amount, percentage
            expected_data = [
                ('digital advertising', '$180,000', '36%'),
                ('content production', '$95,000', '19%'),
                ('events and conferences', '$120,000', '24%'),
                ('marketing tools and software', '$45,000', '9%'),
                ('agency retainer', '$60,000', '12%'),
            ]

            data_rows = budget_table.rows[1:-1]  # Skip header (first) and total (last) rows
            matched = 0
            for exp_name, exp_amount, exp_pct in expected_data:
                for row in data_rows:
                    row_texts = [c.text.strip().lower() for c in row.cells]
                    if exp_name in row_texts[0]:
                        # Check amount and percentage
                        amount_match = exp_amount.replace(',', '').replace('$', '') in row.cells[1].text.replace(',', '').replace('$', '')
                        pct_match = exp_pct.replace('%', '') in row.cells[2].text.replace('%', '')
                        if amount_match and pct_match:
                            matched += 1
                        break

            if matched == 5:
                print(f"PASS: Component 2 — All 5 data rows have correct budget items, amounts and percentages (0.20 pts)")
                total_score += 0.20
            elif matched >= 3:
                partial = round(0.20 * matched / 5, 2)
                print(f"PARTIAL: Component 2 — {matched}/5 data rows correct ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — Only {matched}/5 data rows correct, expected all 5")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Total row is bold with $500,000 (0.15 points)
    try:
        tables = doc.tables
        budget_table = None
        for table in tables:
            if len(table.rows) >= 7 and len(table.columns) == 3:
                header_texts = [c.text.strip().lower() for c in table.rows[0].cells]
                if any('line item' in h or 'budget' in h or '% of total' in h for h in header_texts):
                    budget_table = table
                    break

        if budget_table is None:
            print(f"FAIL: Component 3 — No budget table found for total row validation")
        else:
            total_row = budget_table.rows[-1]
            total_label = total_row.cells[0].text.strip().lower()
            total_amount = total_row.cells[1].text.strip()

            if 'total' not in total_label:
                print(f"FAIL: Component 3 — Last row label is {total_label!r}, expected 'Total'")
            else:
                # Check amount is $500,000
                amount_clean = total_amount.replace(',', '').replace('$', '').strip()
                if '500000' not in amount_clean and '500,000' not in total_amount:
                    print(f"FAIL: Component 3 — Total amount is {total_amount!r}, expected '$500,000'")
                else:
                    # Check bold formatting on total row cells
                    # Ground truth: Total row should be bold
                    bold_found = False
                    for cell in total_row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.bold or run.font.bold:
                                    bold_found = True
                                    break
                        if bold_found:
                            break

                    if bold_found:
                        print(f"PASS: Component 3 — Total row found with $500,000 and bold formatting (0.15 pts)")
                        total_score += 0.15
                    else:
                        # Amount correct but bold not confirmed — give partial credit
                        print(f"PARTIAL: Component 3 — Total row has $500,000 but bold formatting not detected (0.08 pts)")
                        total_score += 0.08
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 4 numbered list items with correct action items (0.20 points)
    # Expected: 'List Number' style paragraphs containing the 4 action items
    try:
        list_number_paras = [p for p in doc.paragraphs if p.style.name == 'List Number']

        expected_action_items = [
            'secure vendor contracts by march 15',
            'launch spring campaign by april 1',
            'complete mid-year review by july 15',
            'submit q4 forecast by october 1',
        ]

        if len(list_number_paras) < 4:
            print(f"FAIL: Component 4 — Only {len(list_number_paras)} numbered list items found, expected 4")
        else:
            matched_actions = 0
            for expected in expected_action_items:
                for para in list_number_paras:
                    para_lower = para.text.strip().lower()
                    # Check for key words in each action item
                    exp_words = expected.split()[:3]  # First 3 words as key identifier
                    if all(w in para_lower for w in exp_words):
                        matched_actions += 1
                        break

            if matched_actions == 4:
                print(f"PASS: Component 4 — All 4 numbered list action items found with correct content (0.20 pts)")
                total_score += 0.20
            elif matched_actions >= 2:
                partial = round(0.20 * matched_actions / 4, 2)
                print(f"PARTIAL: Component 4 — {matched_actions}/4 numbered list items matched ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Only {matched_actions}/4 action items found as numbered list items")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: "Budget Summary" and "Key Action Items" headings are bold and 14pt (0.10 points)
    # Both headings should be bold at 14pt per ground truth spec
    try:
        found_budget_summary = False
        found_key_action_items = False

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text == 'Budget Summary':
                # Check bold and 14pt on runs
                for run in para.runs:
                    run_bold = run.bold or run.font.bold
                    run_size = run.font.size
                    size_ok = (run_size == Pt(14)) if run_size is not None else False
                    if run_bold and size_ok:
                        found_budget_summary = True
                        break
            elif para_text == 'Key Action Items':
                for run in para.runs:
                    run_bold = run.bold or run.font.bold
                    run_size = run.font.size
                    size_ok = (run_size == Pt(14)) if run_size is not None else False
                    if run_bold and size_ok:
                        found_key_action_items = True
                        break

        if found_budget_summary and found_key_action_items:
            print(f"PASS: Component 5 — Both 'Budget Summary' and 'Key Action Items' headings are bold 14pt (0.10 pts)")
            total_score += 0.10
        elif found_budget_summary or found_key_action_items:
            found_which = 'Budget Summary' if found_budget_summary else 'Key Action Items'
            print(f"PARTIAL: Component 5 — Only {found_which!r} found as bold 14pt heading (0.05 pts)")
            total_score += 0.05
        else:
            # Check if headings exist at all without strict formatting check
            budget_exists = any(p.text.strip() == 'Budget Summary' for p in doc.paragraphs)
            action_exists = any(p.text.strip() == 'Key Action Items' for p in doc.paragraphs)
            if budget_exists and action_exists:
                print(f"PARTIAL: Component 5 — Both headings exist but formatting (bold/14pt) not confirmed (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 5 — budget_summary={budget_exists}, key_action_items={action_exists}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
