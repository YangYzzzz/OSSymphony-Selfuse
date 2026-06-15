"""
Initial Setup: Marketing Budget Justification Document (narrative form)
Task ID: writer_mktg_027
Domain: libreoffice_writer

Creates a 3-page narrative document with budget items and action items
embedded in flowing paragraph text (NOT in a table or numbered list).
The agent's task is to restructure this into a table + numbered list.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_027'
OUTPUT = f'{WORKDIR}/budget_justification_2026.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page 1: Executive Overview ---
    heading = doc.add_heading('Marketing Budget Justification 2026', level=1)
    heading.runs[0].font.size = Pt(16)

    doc.add_heading('Executive Overview', level=2)

    para = doc.add_paragraph(
        'The following document provides a comprehensive justification for the proposed '
        '2026 marketing budget allocation. As we enter a critical growth phase, the marketing '
        'department has carefully evaluated all expenditures to ensure maximum return on '
        'investment and alignment with corporate strategic objectives.'
    )

    doc.add_paragraph(
        'This year, our total proposed budget stands at five hundred thousand dollars, '
        'distributed across five major spending categories. Each category has been '
        'scrutinized by department leadership and benchmarked against industry standards '
        'to confirm its necessity and projected impact on revenue generation.'
    )

    doc.add_paragraph(
        'The marketing landscape has evolved significantly since our last budget cycle. '
        'Digital channels now account for the majority of consumer touchpoints, and our '
        'investment strategy reflects this shift. We have maintained investments in proven '
        'areas while introducing new capabilities to capture emerging market opportunities '
        'in the Asia-Pacific and European regions.'
    )

    doc.add_page_break()

    # --- Page 2: Budget Breakdown (narrative — no table) ---
    doc.add_heading('Budget Breakdown and Rationale', level=2)

    doc.add_paragraph(
        'The single largest investment in our plan is Digital advertising ($180,000), '
        'which represents our programmatic, paid search, and social media campaigns. '
        'This allocation reflects a 15% increase over last year due to expanded campaign '
        'footprint and higher CPM rates in key demographic segments. Analytics data from '
        'Q3 2025 confirms a 4.2x ROAS on paid channels, justifying the increase.'
    )

    doc.add_paragraph(
        'Content production ($95,000) covers video creation, copywriting, graphic design, '
        'and photography for all brand and product materials. Our in-house team handles '
        'approximately 60% of production volume, with the remainder outsourced to vetted '
        'creative agencies. This allocation includes a new investment in interactive '
        'content formats such as quizzes and configurators to improve lead quality.'
    )

    doc.add_paragraph(
        'Events and conferences ($120,000) funds our presence at eight industry trade shows '
        'and two proprietary customer events. Trade show participation has consistently '
        'generated qualified pipeline exceeding 3x the investment cost. The two owned '
        'events, launched in 2024, generated $2.3M in influenced revenue last year, '
        'making this one of our most cost-efficient budget categories.'
    )

    doc.add_paragraph(
        'Marketing tools and software ($45,000) encompasses our marketing automation '
        'platform, CRM integrations, analytics subscriptions, and design software licenses. '
        'This category has remained stable year-over-year, with modest increases tied to '
        'user seat expansions and new capability modules required by the data team. '
        'Consolidation of three legacy tools in H2 2025 will offset approximately $8,000 '
        'in new license costs.'
    )

    doc.add_paragraph(
        'The Agency retainer ($60,000) covers our relationship with Brandline Partners, '
        'our strategic communications agency of record since 2023. This retainer includes '
        'monthly strategy sessions, PR outreach, media relations, and crisis communications '
        'support. An independent review conducted in November 2025 confirmed the agency '
        'delivers approximately 2.5x the value of comparable in-house capabilities.'
    )

    doc.add_page_break()

    # --- Page 3: Execution Plan and Action Items ---
    doc.add_heading('Execution Plan and Next Steps', level=2)

    doc.add_paragraph(
        'Approval of the 2026 marketing budget requires coordination across finance, '
        'procurement, and executive leadership. The following timeline has been established '
        'to ensure all commitments are in place prior to the start of Q2 campaign activity.'
    )

    doc.add_paragraph(
        'The first priority is to Secure vendor contracts by March 15. All supplier '
        'agreements, including agency retainer renewals and platform subscriptions, must be '
        'executed no later than the 15th to avoid service interruptions and lock in '
        'negotiated pricing. Legal review of the Brandline Partners agreement is already '
        'underway.'
    )

    doc.add_paragraph(
        'Following contract execution, the team will Launch spring campaign by April 1. '
        'The spring campaign represents our largest single initiative of the year, targeting '
        'enterprise accounts in the financial services and healthcare verticals. Creative '
        'assets are in final review and media placements have been pre-negotiated pending '
        'budget approval.'
    )

    doc.add_paragraph(
        'At the midpoint of the year, we will Complete mid-year review by July 15. This '
        'review will assess actual spend versus budget, evaluate campaign performance '
        'against KPIs, and make reallocation recommendations for H2 if market conditions '
        'warrant. The review findings will be presented to the CMO and CFO in a joint '
        'session no later than July 30.'
    )

    doc.add_paragraph(
        'As the fiscal year draws to a close, the team will Submit Q4 forecast by October 1. '
        'This forecast will inform the preliminary 2027 budget discussions and provide '
        'finance with accurate accrual projections for year-end close. Historical accuracy '
        'of our Q4 forecasts has been within 5% of actuals for three consecutive years.'
    )

    doc.add_paragraph(
        'We appreciate the leadership team\'s continued support of the marketing function '
        'and remain committed to delivering measurable business impact from every dollar '
        'invested. Questions regarding this justification document should be directed to '
        'the Director of Marketing Operations, Alexandra Kim, at ext. 4872.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
