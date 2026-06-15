"""
Reward Script: Set document title to bold, 18pt, center-aligned
Task ID: writer_tech_005
Domain: libreoffice_writer
Scoring:
  Component 1 — Title bold (0.35)
  Component 2 — Title 18pt font size (0.35)
  Component 3 — Title center-aligned (0.30)
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_005'


def persist_app_state(domain: str):
    """Best-effort save of any unsaved GUI edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one paragraph with the title text
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    title_para = doc.paragraphs[0]
    title_text = title_para.text.strip()
    if 'CloudSync API Documentation v2.1' not in title_text:
        print(f"FAIL: First paragraph is not the expected title. Found: {repr(title_text[:80])}")
        print("REWARD: 0.0")
        return 0.0

    print(f"PRECONDITION: Title paragraph found: {repr(title_text[:60])}")

    # Component 1: Title text is bold (0.35 points)
    # All runs in the title paragraph must be bold
    try:
        runs = title_para.runs
        if len(runs) == 0:
            print("FAIL: Component 1 — Title paragraph has no runs")
        else:
            all_bold = all(run.font.bold is True for run in runs if run.text.strip())
            if all_bold:
                print(f"PASS: Component 1 — Title is bold (0.35 pts)")
                total_score += 0.35
            else:
                bold_values = [(run.text[:20], run.font.bold) for run in runs]
                print(f"FAIL: Component 1 — Expected all runs bold=True, found: {bold_values}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Title font size is 18pt (0.35 points)
    try:
        runs = title_para.runs
        if len(runs) == 0:
            print("FAIL: Component 2 — Title paragraph has no runs")
        else:
            all_18pt = all(
                run.font.size is not None and abs(run.font.size.pt - 18.0) < 0.5
                for run in runs if run.text.strip()
            )
            if all_18pt:
                actual_pts = [run.font.size.pt if run.font.size else None for run in runs]
                print(f"PASS: Component 2 — Title font size is 18pt (values: {actual_pts}) (0.35 pts)")
                total_score += 0.35
            else:
                actual_pts = [run.font.size.pt if run.font.size else None for run in runs]
                print(f"FAIL: Component 2 — Expected 18pt, found: {actual_pts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Title paragraph is center-aligned (0.30 points)
    try:
        alignment = title_para.paragraph_format.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            print(f"PASS: Component 3 — Title is center-aligned (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 3 — Expected CENTER alignment, found: {alignment}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved GUI state before verification
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
