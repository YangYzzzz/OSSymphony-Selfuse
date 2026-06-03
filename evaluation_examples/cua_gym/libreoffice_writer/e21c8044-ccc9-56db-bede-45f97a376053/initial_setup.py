"""
Initial Setup: Create a technical document about HTTP/1.1 that references RFC 7231
but has no formal bibliography entries, citation marks, or bibliography table.
Task ID: writer_tech_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('HTTP/1.1 Semantics and Content: A Technical Overview', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author Info ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Prepared by: Elena Vasquez, Senior Network Engineer')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Last Updated: March 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)

    p = doc.add_paragraph()
    p.add_run('The Hypertext Transfer Protocol (HTTP) is the foundation of data communication on the World Wide Web. ')
    p.add_run('This document provides a technical overview of the semantics and content aspects of HTTP/1.1 as defined in ')
    run_ref = p.add_run('RFC 7231 - HTTP/1.1 Semantics and Content')
    run_ref.font.italic = True
    p.add_run('. The specification, authored by R. Fielding and J. Reschke and published by the Internet Engineering Task Force (IETF) in June 2014, ')
    p.add_run('supersedes the earlier RFC 2616 and establishes the current standard for HTTP message semantics.')

    p2 = doc.add_paragraph()
    p2.add_run('HTTP/1.1 remains widely deployed across the internet despite the adoption of HTTP/2 and HTTP/3. ')
    p2.add_run('Understanding the semantic foundations laid out in RFC 7231 is essential for web developers, ')
    p2.add_run('API designers, and network engineers working with modern web infrastructure.')

    # --- Section 2: Request Methods ---
    doc.add_heading('2. Request Methods', level=1)

    p3 = doc.add_paragraph()
    p3.add_run('RFC 7231 defines the semantics for the standard HTTP request methods. ')
    p3.add_run('Each method indicates the desired action to be performed on the identified resource. ')
    p3.add_run('The following table summarizes the core methods:')

    # Methods table
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'

    headers = ['Method', 'Safe', 'Idempotent']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    methods_data = [
        ['GET', 'Yes', 'Yes'],
        ['HEAD', 'Yes', 'Yes'],
        ['POST', 'No', 'No'],
        ['PUT', 'No', 'Yes'],
        ['DELETE', 'No', 'Yes'],
        ['CONNECT', 'No', 'No'],
        ['OPTIONS', 'Yes', 'Yes'],
        ['TRACE', 'Yes', 'Yes'],
    ]
    for r, row_data in enumerate(methods_data, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.add_run('The distinction between safe and idempotent methods, as outlined in RFC 7231 Section 4.2, ')
    p4.add_run('has significant implications for caching, retry logic, and proxy behavior. ')
    p4.add_run('Safe methods are those that do not modify server state, while idempotent methods ')
    p4.add_run('produce the same result regardless of how many times they are executed.')

    # --- Section 3: Status Codes ---
    doc.add_heading('3. Response Status Codes', level=1)

    p5 = doc.add_paragraph()
    p5.add_run('The status code of a response is a three-digit integer code giving the result of the attempt ')
    p5.add_run('to understand and satisfy the request. RFC 7231 organizes status codes into five classes:')

    status_classes = [
        ('1xx (Informational)', 'The request was received, continuing process.'),
        ('2xx (Successful)', 'The request was successfully received, understood, and accepted.'),
        ('3xx (Redirection)', 'Further action needs to be taken to complete the request.'),
        ('4xx (Client Error)', 'The request contains bad syntax or cannot be fulfilled.'),
        ('5xx (Server Error)', 'The server failed to fulfill an apparently valid request.'),
    ]
    for code_class, description in status_classes:
        p = doc.add_paragraph()
        run_code = p.add_run(code_class + ': ')
        run_code.bold = True
        run_code.font.size = Pt(10)
        run_desc = p.add_run(description)
        run_desc.font.size = Pt(10)

    p6 = doc.add_paragraph()
    p6.add_run('Notable additions in RFC 7231 include clarifications on the use of 200 OK with different methods, ')
    p6.add_run('the semantics of 203 Non-Authoritative Information, and improved guidance on ')
    p6.add_run('when to use 301 Moved Permanently versus 308 Permanent Redirect.')

    # --- Section 4: Content Negotiation ---
    doc.add_heading('4. Content Negotiation', level=1)

    p7 = doc.add_paragraph()
    p7.add_run('One of the key contributions of RFC 7231 is the formalization of content negotiation mechanisms. ')
    p7.add_run('Content negotiation allows clients and servers to agree on the most appropriate representation ')
    p7.add_run('of a resource based on capabilities and preferences.')

    doc.add_heading('4.1 Proactive Negotiation', level=2)

    p8 = doc.add_paragraph()
    p8.add_run('In proactive (server-driven) negotiation, the server selects the representation based on ')
    p8.add_run('request header fields such as Accept, Accept-Language, Accept-Encoding, and Accept-Charset. ')
    p8.add_run('The server evaluates these preferences against available representations and responds with ')
    p8.add_run('the best match. This approach is described in Section 5.3 of RFC 7231.')

    doc.add_heading('4.2 Reactive Negotiation', level=2)

    p9 = doc.add_paragraph()
    p9.add_run('Reactive (agent-driven) negotiation occurs when the server provides a list of available ')
    p9.add_run('representations, typically via a 300 Multiple Choices or 406 Not Acceptable response, ')
    p9.add_run('allowing the user agent to select the preferred variant.')

    # --- Section 5: Header Fields ---
    doc.add_heading('5. Representation Header Fields', level=1)

    p10 = doc.add_paragraph()
    p10.add_run('RFC 7231 defines several important header fields for content representation:')

    header_fields = [
        ('Content-Type', 'Indicates the media type of the representation, e.g., text/html or application/json.'),
        ('Content-Encoding', 'Specifies the encoding transformations applied, such as gzip or deflate compression.'),
        ('Content-Language', 'Describes the natural language(s) of the intended audience.'),
        ('Content-Location', 'Provides a URI reference for a resource corresponding to the representation.'),
    ]
    for field_name, field_desc in header_fields:
        p = doc.add_paragraph(style='List Bullet')
        run_name = p.add_run(field_name + ': ')
        run_name.bold = True
        p.add_run(field_desc)

    # --- Section 6: Practical Implications ---
    doc.add_heading('6. Practical Implications for Modern Development', level=1)

    p11 = doc.add_paragraph()
    p11.add_run('Understanding RFC 7231 is critical for several areas of modern software development:')

    implications = [
        'RESTful API design relies heavily on the method semantics defined in RFC 7231 to ensure consistent behavior across implementations.',
        'Caching strategies depend on understanding safe and idempotent methods, as well as response freshness indicators.',
        'Security considerations, such as those related to the CONNECT method and sensitive header fields, are addressed in Section 9.',
        'Proxy and gateway implementations must correctly handle method semantics to maintain protocol conformance.',
        'Web framework developers reference RFC 7231 to ensure correct default status codes and content type handling.',
    ]
    for imp in implications:
        doc.add_paragraph(imp, style='List Number')

    # --- Conclusion ---
    doc.add_heading('7. Conclusion', level=1)

    p12 = doc.add_paragraph()
    p12.add_run('RFC 7231 remains a cornerstone document for understanding HTTP/1.1 semantics. ')
    p12.add_run('Its precise definitions of request methods, status codes, and content negotiation ')
    p12.add_run('mechanisms continue to guide the development of web standards and applications. ')
    p12.add_run('While newer protocols like HTTP/2 (RFC 7540) and HTTP/3 (RFC 9114) have been developed, ')
    p12.add_run('the semantic foundations established in RFC 7231 remain largely unchanged and continue ')
    p12.add_run('to serve as the authoritative reference for HTTP message semantics.')

    doc.add_paragraph()

    # --- References section (informal, no bibliography database) ---
    doc.add_heading('References', level=1)

    refs = [
        'RFC 7231 - Hypertext Transfer Protocol (HTTP/1.1): Semantics and Content, R. Fielding, J. Reschke, IETF, June 2014',
        'RFC 7230 - Hypertext Transfer Protocol (HTTP/1.1): Message Syntax and Routing',
        'RFC 7232 - Hypertext Transfer Protocol (HTTP/1.1): Conditional Requests',
        'RFC 7234 - Hypertext Transfer Protocol (HTTP/1.1): Caching',
        'RFC 7235 - Hypertext Transfer Protocol (HTTP/1.1): Authentication',
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Number')
        p.add_run(ref)

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
