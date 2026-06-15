"""
Initial Setup: Mail merge data source with 12 records in LibreOffice Writer
Task ID: writer_mt_015
Domain: libreoffice_writer
"""

import csv
import os
import shlex
import subprocess
import time
import xml.etree.ElementTree as ET

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_015'
DOC_OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
CSV_PATH = f'{WORKDIR}/TeamMembers.csv'

# Data source registration paths
LO_PROFILE = os.path.expanduser('~/.config/libreoffice/4/user')
DB_DIR = f'{LO_PROFILE}/database'
DATASOURCE_DIR = f'{LO_PROFILE}/datasources'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# 12 realistic team member records
TEAM_DATA = [
    ['FirstName', 'LastName', 'Email', 'Department'],
    ['Sarah', 'Chen', 'schen@company.com', 'Engineering'],
    ['Marcus', 'Johnson', 'mjohnson@company.com', 'Marketing'],
    ['Priya', 'Patel', 'ppatel@company.com', 'Finance'],
    ['David', 'Kim', 'dkim@company.com', 'Engineering'],
    ['Elena', 'Rodriguez', 'erodriguez@company.com', 'Human Resources'],
    ['Thomas', 'Anderson', 'tanderson@company.com', 'Sales'],
    ['Aisha', 'Mohammed', 'amohammed@company.com', 'Engineering'],
    ['Robert', 'Taylor', 'rtaylor@company.com', 'Marketing'],
    ['Mei', 'Wong', 'mwong@company.com', 'Finance'],
    ['Carlos', 'Garcia', 'cgarcia@company.com', 'Sales'],
    ['Jennifer', 'Smith', 'jsmith@company.com', 'Human Resources'],
    ['Alex', 'Nakamura', 'anakamura@company.com', 'Engineering'],
]


def create_csv():
    """Create the TeamMembers CSV data source with 12 records."""
    with open(CSV_PATH, 'w', newline='') as f:
        writer = csv.writer(f)
        for row in TEAM_DATA:
            writer.writerow(row)
    print(f'CSV data source created: {CSV_PATH} ({len(TEAM_DATA) - 1} records)')


def create_document():
    """Create a mail merge template document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    heading = doc.add_heading('Team Communication Template', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction paragraph
    doc.add_paragraph(
        'This document serves as a mail merge template for sending personalized '
        'communications to team members. The data source "TeamMembers" contains '
        'current staff records organized by department.'
    )

    doc.add_paragraph('')

    # Mail merge fields placeholder text
    p = doc.add_paragraph()
    run = p.add_run('Dear ')
    run.font.size = Pt(12)
    run = p.add_run('[FirstName]')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(' ')
    run.font.size = Pt(12)
    run = p.add_run('[LastName]')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(',')
    run.font.size = Pt(12)

    doc.add_paragraph(
        'We are writing to inform you about upcoming changes in the '
        '[Department] department. Please review the attached materials '
        'and respond to [Email] with any questions.'
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'This template is connected to the TeamMembers data source. '
        'Use the data source browser (F4) to view and manage records.'
    )

    # Closing
    doc.add_paragraph('')
    p = doc.add_paragraph('Best regards,')
    doc.add_paragraph('HR Communications Team')

    doc.save(DOC_OUTPUT)
    print(f'Document created: {DOC_OUTPUT}')


def register_datasource():
    """Register the CSV as a LibreOffice data source named TeamMembers."""
    os.makedirs(DB_DIR, exist_ok=True)

    # Create a .odb database file that references the CSV directory
    # LibreOffice uses a flat file driver for CSV data sources
    # We register through the datasources configuration

    # Method: Create an .odb file pointing to the CSV
    odb_path = f'{DB_DIR}/TeamMembers.odb'

    # Use LibreOffice's built-in SDBC flat file driver approach
    # Register the data source via xcu configuration
    os.makedirs(DATASOURCE_DIR, exist_ok=True)

    # Create the data source registration in LibreOffice config
    registrations_dir = f'{LO_PROFILE}/registrymodifications.xcu'

    # Create a simple .odb (Base database) that points to the CSV directory
    # Actually, the easiest reliable method is to use a macro or the registrations

    # Write a registrations.xcu that maps "TeamMembers" to a flat-file SDBC URL
    # pointing to the CSV directory

    # First, let's create an sdbc connection string for flat CSV files
    # The LibreOffice flat file driver URL format is:
    # sdbc:flat:<directory_path>

    # We need to add the registration to the LO config
    reg_path = f'{LO_PROFILE}/registrymodifications.xcu'

    # Read existing file if present
    if os.path.exists(reg_path):
        with open(reg_path, 'r') as f:
            content = f.read()
    else:
        content = '''<?xml version="1.0" encoding="UTF-8"?>
<oor:items xmlns:oor="http://openoffice.org/2001/registry"
           xmlns:xs="http://www.w3.org/2001/XMLSchema"
           xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
</oor:items>'''

    # Add data source registration entries if not already there
    if 'TeamMembers' not in content:
        # Insert before closing </oor:items>
        registration_entries = f'''
<item oor:path="/org.openoffice.Office.DataAccess/RegisteredNames">
  <node oor:name="TeamMembers" oor:op="replace">
    <prop oor:name="Location" oor:type="xs:string">
      <value>sdbc:flat:{WORKDIR}</value>
    </prop>
    <prop oor:name="Name" oor:type="xs:string">
      <value>TeamMembers</value>
    </prop>
  </node>
</item>'''
        content = content.replace('</oor:items>', registration_entries + '\n</oor:items>')

        with open(reg_path, 'w') as f:
            f.write(content)
        print(f'Data source "TeamMembers" registered in LibreOffice config')
    else:
        print('Data source "TeamMembers" already registered')


def create_initial():
    create_csv()
    create_document()
    register_datasource()

    # Kill any existing LibreOffice instances first
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(2)

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=5.0)

    # Send F4 to open the data source browser via xdg-utils or python-xlib
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    time.sleep(2)

    # Try to send F4 using python3 with Xlib (ctypes approach)
    f4_script = '''
import ctypes, ctypes.util, time
try:
    xlib = ctypes.cdll.LoadLibrary(ctypes.util.find_library("X11"))
    xtst = ctypes.cdll.LoadLibrary(ctypes.util.find_library("Xtst"))
    display = xlib.XOpenDisplay(None)
    if display:
        # F4 keycode is typically 70
        xtst.XTestFakeKeyEvent(display, 70, True, 0)   # press
        xtst.XTestFakeKeyEvent(display, 70, False, 0)   # release
        xlib.XFlush(display)
        xlib.XCloseDisplay(display)
        print("F4 sent via XTest")
    else:
        print("Could not open display")
except Exception as e:
    print(f"F4 send failed: {e}")
'''
    subprocess.Popen(
        ['python3', '-c', f4_script],
        env=env,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    time.sleep(2)

    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0, data source browser toggled (F4)')


create_initial()
