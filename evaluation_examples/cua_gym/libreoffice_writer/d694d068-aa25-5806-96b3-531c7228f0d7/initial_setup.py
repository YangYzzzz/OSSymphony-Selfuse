"""
Initial Setup: Create a technical document without a glossary section.
Task ID: writer_tech_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('CloudSync Platform — Technical Specification', level=0)

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'CloudSync is a distributed file synchronization platform designed to '
        'provide seamless real-time collaboration across teams. This document '
        'outlines the core architecture, design decisions, and implementation '
        'guidelines for the engineering team.'
    )
    doc.add_paragraph(
        'The platform supports concurrent editing, conflict resolution, and '
        'automatic versioning. It is built on a microservices architecture with '
        'event-driven communication between components.'
    )

    # --- Section 2: System Architecture ---
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        'The system is composed of five primary services: the Authentication '
        'Gateway, File Storage Service, Sync Engine, Notification Service, and '
        'the Analytics Pipeline. Each service communicates through a message '
        'broker using asynchronous event streams.'
    )

    doc.add_heading('2.1 Authentication Gateway', level=2)
    doc.add_paragraph(
        'The Authentication Gateway handles user identity verification, token '
        'issuance, and session management. It supports multiple identity providers '
        'including Google Workspace, Microsoft Entra ID, and SAML-based enterprise '
        'SSO configurations. Rate limiting is enforced at 500 requests per minute '
        'per client.'
    )

    doc.add_heading('2.2 File Storage Service', level=2)
    doc.add_paragraph(
        'Files are stored in an object storage backend with content-addressable '
        'hashing. The service maintains metadata in a PostgreSQL database and '
        'supports chunked uploads for files larger than 100 MB. Each chunk is '
        'individually checksummed to ensure data integrity during transfer.'
    )

    doc.add_heading('2.3 Sync Engine', level=2)
    doc.add_paragraph(
        'The Sync Engine uses operational transformation to merge concurrent edits. '
        'Conflict detection is based on vector clocks, and resolution follows a '
        'last-writer-wins strategy for non-collaborative documents. For shared '
        'editing sessions, changes are streamed via WebSocket connections with a '
        'target latency of under 200 milliseconds.'
    )

    # --- Section 3: Data Model ---
    doc.add_heading('3. Data Model', level=1)
    doc.add_paragraph(
        'The core entities in the data model are Users, Workspaces, Files, and '
        'Versions. A workspace contains a hierarchical folder structure, and each '
        'file maintains a complete version history. Soft deletion is used '
        'throughout, with a 30-day retention window before permanent purging.'
    )

    # Add a simple table for the data model
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Entity', 'Primary Key', 'Description']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    data = [
        ['User', 'user_id (UUID)', 'Registered platform user with role assignments'],
        ['Workspace', 'workspace_id (UUID)', 'Shared collaboration space with access controls'],
        ['File', 'file_id (UUID)', 'Document or binary asset within a workspace'],
        ['Version', 'version_id (UUID)', 'Immutable snapshot of a file at a point in time'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Section 4: Deployment ---
    doc.add_heading('4. Deployment Guidelines', level=1)
    doc.add_paragraph(
        'All services are containerized using Docker and orchestrated with '
        'Kubernetes. The production cluster runs on three availability zones with '
        'automatic failover. CI/CD pipelines are configured in GitHub Actions with '
        'staging and production environments. Blue-green deployments are used for '
        'zero-downtime releases.'
    )

    doc.add_heading('4.1 Environment Variables', level=2)
    doc.add_paragraph(
        'Each service reads configuration from environment variables injected via '
        'Kubernetes ConfigMaps and Secrets. Database connection strings, external '
        'service endpoints, and feature flags are all managed through this '
        'mechanism. No credentials are stored in source control.'
    )

    # --- Section 5: Security Considerations ---
    doc.add_heading('5. Security Considerations', level=1)
    doc.add_paragraph(
        'All inter-service communication is encrypted using mutual TLS. User data '
        'at rest is encrypted with AES-256. Access tokens expire after 15 minutes '
        'and refresh tokens after 7 days. Audit logs are retained for 90 days and '
        'shipped to a centralized SIEM for anomaly detection.'
    )
    doc.add_paragraph(
        'Penetration testing is conducted quarterly by an external security firm. '
        'Vulnerability disclosures follow a responsible disclosure policy with a '
        '90-day remediation window.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
