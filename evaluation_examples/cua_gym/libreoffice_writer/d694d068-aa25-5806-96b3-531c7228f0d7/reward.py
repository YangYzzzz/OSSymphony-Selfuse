"""
Reward Script: Create a glossary table at the end of the document
Task ID: writer_tech_032
Domain: libreoffice_writer
Scoring:
  Component 1 (0.2): Glossary heading exists with Heading 1 style
  Component 2 (0.2): Glossary table with 2-column structure (Term, Definition header)
  Component 3 (0.3): All 5 required terms present (API, SDK, REST, JSON, OAuth)
  Component 4 (0.3): Each term has a meaningful definition (>20 chars)
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_032'

REQUIRED_TERMS = ['API', 'SDK', 'REST', 'JSON', 'OAuth']


def persist_app_state(domain):
    """Best-effort save in case document is open in LibreOffice."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Glossary heading exists with Heading 1 style (0.2 points)
    # The task requires a 'Glossary' heading at the end of the document.
    # This must be a NEW heading not present in the initial document.
    try:
        glossary_heading_found = False
        glossary_heading_idx = -1
        for i, para in enumerate(doc.paragraphs):
            if 'heading' in para.style.name.lower() and 'glossary' in para.text.lower():
                glossary_heading_found = True
                glossary_heading_idx = i
                break

        if glossary_heading_found:
            # Verify it is at or near the end of the document
            # (there should be no more content paragraphs with substantial text after it,
            # aside from possibly empty paragraphs)
            is_near_end = True
            for j in range(glossary_heading_idx + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip() and 'heading' in doc.paragraphs[j].style.name.lower():
                    # Another heading after Glossary means it's not at the end
                    is_near_end = False
                    break

            if is_near_end:
                print(f"PASS: Component 1 — Glossary heading found at para {glossary_heading_idx} with style '{doc.paragraphs[glossary_heading_idx].style.name}' (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 1 — Glossary heading found but not at end of document")
        else:
            print(f"FAIL: Component 1 — No Glossary heading found with a Heading style")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Glossary table with correct structure (0.2 points)
    # Must have a 2-column table with header row containing 'Term' and 'Definition'
    try:
        glossary_table = None
        # Look for a table with 2 columns and 'Term'/'Definition' headers
        for table in doc.tables:
            if len(table.columns) == 2:
                header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'term' in header_cells and 'definition' in header_cells:
                    glossary_table = table
                    break

        if glossary_table is not None:
            num_data_rows = len(glossary_table.rows) - 1  # exclude header
            if num_data_rows >= 5:
                print(f"PASS: Component 2 — Glossary table found: 2 columns, {num_data_rows} data rows (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 2 — Glossary table found but only {num_data_rows} data rows (need >= 5)")
        else:
            print(f"FAIL: Component 2 — No 2-column table with Term/Definition headers found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Components 3 & 4 depend on finding the glossary table
    if glossary_table is None:
        print(f"FAIL: Components 3 & 4 — No glossary table to check terms/definitions")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Build a dict of term -> definition from the table
    term_def_map = {}
    for row_idx in range(1, len(glossary_table.rows)):
        row = glossary_table.rows[row_idx]
        term_text = row.cells[0].text.strip()
        def_text = row.cells[1].text.strip()
        term_def_map[term_text.upper()] = def_text

    # Component 3: All 5 required terms present (0.3 points, 0.06 each)
    try:
        terms_found = 0
        for term in REQUIRED_TERMS:
            if term.upper() in term_def_map:
                terms_found += 1
                print(f"  PASS: Term '{term}' present in glossary")
            else:
                print(f"  FAIL: Term '{term}' missing from glossary")

        comp3_score = round(terms_found * 0.06, 2)
        if terms_found == len(REQUIRED_TERMS):
            print(f"PASS: Component 3 — All {terms_found}/{len(REQUIRED_TERMS)} terms found ({comp3_score} pts)")
        else:
            print(f"PARTIAL: Component 3 — {terms_found}/{len(REQUIRED_TERMS)} terms found ({comp3_score} pts)")
        total_score += comp3_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Each term has a meaningful definition (0.3 points, 0.06 each)
    # A meaningful definition should be at least 20 characters
    try:
        defs_ok = 0
        for term in REQUIRED_TERMS:
            definition = term_def_map.get(term.upper(), "")
            if len(definition) >= 20:
                defs_ok += 1
                print(f"  PASS: '{term}' definition is meaningful ({len(definition)} chars)")
            else:
                print(f"  FAIL: '{term}' definition too short or missing ({len(definition)} chars): '{definition}'")

        comp4_score = round(defs_ok * 0.06, 2)
        if defs_ok == len(REQUIRED_TERMS):
            print(f"PASS: Component 4 — All {defs_ok}/{len(REQUIRED_TERMS)} definitions are meaningful ({comp4_score} pts)")
        else:
            print(f"PARTIAL: Component 4 — {defs_ok}/{len(REQUIRED_TERMS)} definitions are meaningful ({comp4_score} pts)")
        total_score += comp4_score
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
