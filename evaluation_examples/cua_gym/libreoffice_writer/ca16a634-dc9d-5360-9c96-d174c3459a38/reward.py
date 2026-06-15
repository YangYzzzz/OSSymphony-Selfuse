"""
Reward Script: Employee Onboarding Checklist in LibreOffice Writer
Task ID: writer_wf_085
Domain: libreoffice_writer
Scoring:
  Component 1: Title "New Employee Onboarding Checklist" present (0.15)
  Component 2: Employee info table with 5 required fields (0.15)
  Component 3: Five section headings present (0.25)
  Component 4: Checklist tables with 4-column structure (0.20)
  Component 5: ~23 total checklist items with checkboxes (0.25)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_085'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have some content (not blank)
    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        print("FAIL: Document is completely blank — no content found")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title "New Employee Onboarding Checklist" present (0.15 points)
    try:
        title_found = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            if 'new employee onboarding checklist' in text_lower:
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 — Title 'New Employee Onboarding Checklist' found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title 'New Employee Onboarding Checklist' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Employee info table with 5 required fields (0.15 points)
    # The employee info table should have fields: Employee Name, Position, Department, Start Date, Manager
    try:
        required_fields = ['employee name', 'position', 'department', 'start date', 'manager']
        fields_found = 0

        # Search all tables for the employee info fields
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip().lower().rstrip(':')
                    for field in required_fields:
                        if field in cell_text:
                            fields_found += 1
                            break

        # Deduplicate by checking unique matches
        found_fields = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip().lower().rstrip(':')
                    for field in required_fields:
                        if field in cell_text:
                            found_fields.add(field)

        num_found = len(found_fields)
        if num_found >= 5:
            print(f"PASS: Component 2 — All 5 employee info fields found: {found_fields} (0.15 pts)")
            total_score += 0.15
        elif num_found >= 3:
            partial = round(0.15 * num_found / 5, 3)
            print(f"PARTIAL: Component 2 — {num_found}/5 fields found: {found_fields} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {num_found}/5 employee info fields found: {found_fields}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Five section headings present (0.25 points)
    # Expected sections: Before First Day, First Day, First Week, First Month, 90-Day Review
    try:
        expected_sections = [
            'before first day',
            'first day',
            'first week',
            'first month',
            '90-day review',
        ]
        # Collect all paragraph text
        all_para_text = [p.text.strip().lower() for p in doc.paragraphs]
        # Also check table cells for section names (some implementations put them in tables)
        all_text_sources = list(all_para_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text_sources.append(cell.text.strip().lower())

        sections_found = set()
        for section_name in expected_sections:
            for text in all_text_sources:
                if section_name in text:
                    sections_found.add(section_name)
                    break
            # Also check: "first day" should not match "before first day" ambiguously
            # Handle "first day" specifically - it should match but not only via "before first day"

        # Special handling: "first day" might match "before first day", so check specifically
        # We need "first day" as a standalone section separate from "before first day"
        if 'first day' in sections_found:
            # Verify there's a standalone "first day" not just "before first day"
            standalone_found = False
            for text in all_text_sources:
                if 'first day' in text and 'before first day' not in text:
                    standalone_found = True
                    break
            if not standalone_found:
                sections_found.discard('first day')

        num_sections = len(sections_found)
        if num_sections >= 5:
            print(f"PASS: Component 3 — All 5 sections found: {sections_found} (0.25 pts)")
            total_score += 0.25
        elif num_sections >= 3:
            partial = round(0.25 * num_sections / 5, 3)
            print(f"PARTIAL: Component 3 — {num_sections}/5 sections found: {sections_found} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {num_sections}/5 sections found: {sections_found}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Checklist tables with 4-column structure (0.20 points)
    # Each checklist table should have columns: Status/Checkbox, Task Description, Responsible Person, Completion Date
    try:
        valid_checklist_tables = 0
        for table in doc.tables:
            if len(table.columns) >= 4 and len(table.rows) >= 2:
                # Check header row for expected column names
                header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                header_text = ' '.join(header_cells)
                # Look for keywords indicating checklist structure
                has_status = any(kw in header_text for kw in ['status', 'checkbox', '☐', 'check'])
                has_task = any(kw in header_text for kw in ['task', 'description', 'item', 'action'])
                has_responsible = any(kw in header_text for kw in ['responsible', 'person', 'owner', 'assigned'])
                has_date = any(kw in header_text for kw in ['date', 'completion', 'deadline', 'due'])

                if (has_task and has_responsible) or (has_status and has_task):
                    valid_checklist_tables += 1

        if valid_checklist_tables >= 5:
            print(f"PASS: Component 4 — {valid_checklist_tables} valid checklist tables found (0.20 pts)")
            total_score += 0.20
        elif valid_checklist_tables >= 3:
            partial = round(0.20 * valid_checklist_tables / 5, 3)
            print(f"PARTIAL: Component 4 — {valid_checklist_tables}/5 checklist tables found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {valid_checklist_tables} valid checklist tables found (need 5)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: ~23 total checklist items with checkboxes (0.25 points)
    # Count rows with checkbox character (☐) across all checklist tables
    try:
        checkbox_items = 0
        for table in doc.tables:
            if len(table.columns) >= 4 and len(table.rows) >= 2:
                # Skip header row, count data rows with checkboxes
                for row_idx in range(1, len(table.rows)):
                    row_text = ' '.join(cell.text.strip() for cell in table.rows[row_idx].cells)
                    if '☐' in row_text or '□' in row_text or '✓' in row_text or '✗' in row_text:
                        checkbox_items += 1

        # Also count non-checkbox data rows in 4+ column tables as items
        total_data_rows = 0
        for table in doc.tables:
            if len(table.columns) >= 4 and len(table.rows) >= 2:
                total_data_rows += len(table.rows) - 1  # subtract header

        # Use the better count (checkbox items if available, else data rows)
        item_count = max(checkbox_items, total_data_rows)

        if item_count >= 20:
            print(f"PASS: Component 5 — {item_count} checklist items found (checkboxes: {checkbox_items}) (0.25 pts)")
            total_score += 0.25
        elif item_count >= 15:
            partial = round(0.25 * item_count / 23, 3)
            print(f"PARTIAL: Component 5 — {item_count}/23 items found ({partial} pts)")
            total_score += partial
        elif item_count >= 10:
            partial = round(0.25 * item_count / 23, 3)
            print(f"PARTIAL: Component 5 — {item_count}/23 items found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {item_count} items found (need ~23)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
