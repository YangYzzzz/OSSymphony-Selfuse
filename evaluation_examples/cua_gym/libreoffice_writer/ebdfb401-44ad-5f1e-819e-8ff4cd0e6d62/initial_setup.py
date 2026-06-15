"""
Initial Setup: Employment agreement document for comment addition task
Task ID: writer_legal_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('EMPLOYMENT AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Parties ---
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        'This Employment Agreement (the "Agreement") is entered into as of March 15, 2025, '
        'by and between Meridian Technologies, Inc., a Delaware corporation with its principal '
        'offices at 2400 Innovation Drive, Suite 800, San Jose, California 95134 (the "Company"), '
        'and Alexandra R. Vasquez, an individual residing at 1847 Elm Street, Palo Alto, '
        'California 94301 (the "Employee").'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # --- Recitals ---
    doc.add_heading('RECITALS', level=1)

    recitals = [
        'WHEREAS, the Company desires to employ the Employee in the capacity of '
        'Senior Vice President of Product Development, and the Employee desires to '
        'accept such employment on the terms and conditions set forth herein;',
        'WHEREAS, the Employee possesses specialized knowledge, skills, and experience '
        'that are of significant value to the Company;',
        'WHEREAS, the Company and the Employee wish to establish the terms of the '
        'Employee\'s employment, including compensation, benefits, and post-employment obligations;',
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements herein '
        'contained, and for other good and valuable consideration, the receipt and sufficiency '
        'of which are hereby acknowledged, the parties agree as follows:'
    ]
    for r_text in recitals:
        p = doc.add_paragraph()
        run = p.add_run(r_text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 1: Position and Duties ---
    doc.add_heading('Section 1. Position and Duties', level=2)
    s1_paras = [
        '1.1 The Company hereby employs the Employee as Senior Vice President of Product '
        'Development, reporting directly to the Chief Executive Officer. The Employee shall '
        'perform all duties and responsibilities customarily associated with such position.',
        '1.2 The Employee shall devote substantially all of her business time, attention, '
        'and efforts to the performance of her duties hereunder and shall not engage in any '
        'other business activity, whether or not such activity is pursued for gain, profit, '
        'or other pecuniary advantage, without the prior written consent of the Company\'s '
        'Board of Directors.',
        '1.3 The Employee\'s primary place of employment shall be the Company\'s headquarters '
        'in San Jose, California, subject to reasonable travel requirements as necessary for '
        'the performance of her duties.'
    ]
    for text in s1_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 2: Term ---
    doc.add_heading('Section 2. Term of Employment', level=2)
    s2_paras = [
        '2.1 The term of this Agreement shall commence on April 1, 2025 (the "Start Date") '
        'and shall continue until terminated by either party in accordance with Section 6 '
        'of this Agreement.',
        '2.2 The initial period of employment shall be for three (3) years from the Start Date '
        '(the "Initial Term"), subject to renewal upon mutual written agreement of the parties.'
    ]
    for text in s2_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 3: Compensation ---
    doc.add_heading('Section 3. Compensation and Benefits', level=2)
    s3_paras = [
        '3.1 Base Salary. The Company shall pay the Employee an annual base salary of '
        'Three Hundred Fifty Thousand Dollars ($350,000), payable in accordance with the '
        'Company\'s standard payroll practices, subject to applicable tax withholdings.',
        '3.2 Annual Bonus. The Employee shall be eligible for an annual performance bonus '
        'of up to forty percent (40%) of her base salary, based on the achievement of '
        'individual and Company performance targets established by the Board of Directors.',
        '3.3 Equity Compensation. Upon commencement of employment, the Employee shall be '
        'granted a stock option to purchase 150,000 shares of the Company\'s common stock, '
        'vesting over a four-year period with a one-year cliff, in accordance with the '
        'Company\'s 2024 Equity Incentive Plan.',
        '3.4 Benefits. The Employee shall be entitled to participate in all benefit programs '
        'offered to similarly situated executives, including health insurance, dental and '
        'vision coverage, life insurance, and the Company\'s 401(k) retirement plan.'
    ]
    for text in s3_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 4: Confidentiality ---
    doc.add_heading('Section 4. Confidentiality', level=2)
    s4_paras = [
        '4.1 The Employee acknowledges that during the course of her employment, she will '
        'have access to and may develop Confidential Information. "Confidential Information" '
        'shall mean all non-public information, whether written, oral, or electronic, relating '
        'to the Company\'s business, including but not limited to trade secrets, customer lists, '
        'product roadmaps, financial data, marketing strategies, and proprietary technology.',
        '4.2 The Employee agrees to hold all Confidential Information in strict confidence and '
        'not to disclose, publish, or otherwise reveal any Confidential Information to any third '
        'party during or after employment, except as required in the performance of her duties '
        'or as authorized in writing by the Company.',
        '4.3 Upon termination of employment, the Employee shall promptly return all documents, '
        'materials, and property containing Confidential Information to the Company.'
    ]
    for text in s4_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 5: Intellectual Property ---
    doc.add_heading('Section 5. Intellectual Property', level=2)
    s5_paras = [
        '5.1 All inventions, discoveries, improvements, and works of authorship conceived or '
        'developed by the Employee during the term of employment, whether or not patentable, '
        'that relate to the Company\'s business shall be the sole and exclusive property of '
        'the Company ("Work Product").',
        '5.2 The Employee hereby assigns to the Company all right, title, and interest in '
        'and to any Work Product, and agrees to execute any documents and take any actions '
        'reasonably necessary to perfect the Company\'s ownership thereof.'
    ]
    for text in s5_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 6: Termination ---
    doc.add_heading('Section 6. Termination', level=2)
    s6_paras = [
        '6.1 Termination by Company for Cause. The Company may terminate the Employee\'s '
        'employment immediately for Cause. "Cause" shall include: (a) material breach of '
        'this Agreement; (b) willful misconduct or gross negligence; (c) conviction of a felony; '
        'or (d) failure to perform duties after written notice and a thirty (30) day cure period.',
        '6.2 Termination Without Cause. The Company may terminate the Employee\'s employment '
        'without Cause upon sixty (60) days\' prior written notice. In such event, the Employee '
        'shall be entitled to severance equal to twelve (12) months of base salary.',
        '6.3 Resignation. The Employee may resign upon thirty (30) days\' prior written notice '
        'to the Company. Voluntary resignation shall not entitle the Employee to any severance '
        'benefits.'
    ]
    for text in s6_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 7: Non-Compete (contains the target clause) ---
    doc.add_heading('Section 7. Non-Competition and Non-Solicitation', level=2)
    s7_paras = [
        '7.1 Non-Competition. During the term of employment and for a period following '
        'termination as specified below, the Employee shall not, directly or indirectly, '
        'engage in, own, manage, operate, or control any business that competes with the '
        'Company within the geographic areas where the Company conducts business.',
        # This paragraph contains the exact target clause text
        '7.2 Duration. The non-compete period shall extend for twenty-four (24) months '
        'following the date of termination of employment, regardless of the reason for '
        'such termination. During this period, the Employee shall refrain from accepting '
        'employment with any Competing Enterprise as defined in Exhibit A.',
        '7.3 Non-Solicitation. For a period of eighteen (18) months following termination, '
        'the Employee shall not, directly or indirectly, solicit, recruit, or hire any '
        'employee, consultant, or contractor of the Company, or induce any such person to '
        'terminate their relationship with the Company.',
        '7.4 Non-Solicitation of Clients. For a period of eighteen (18) months following '
        'termination, the Employee shall not, directly or indirectly, solicit or attempt to '
        'solicit business from any client or customer of the Company with whom the Employee '
        'had material contact during the last twelve (12) months of employment.',
        '7.5 Reasonableness. The Employee acknowledges that the restrictions set forth in '
        'this Section 7 are reasonable in scope, duration, and geographic extent, and are '
        'necessary to protect the Company\'s legitimate business interests.'
    ]
    for text in s7_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 8: Governing Law ---
    doc.add_heading('Section 8. Governing Law and Dispute Resolution', level=2)
    s8_paras = [
        '8.1 This Agreement shall be governed by and construed in accordance with the laws '
        'of the State of California, without regard to its conflict of laws principles.',
        '8.2 Any dispute arising out of or relating to this Agreement shall be resolved '
        'through binding arbitration in Santa Clara County, California, in accordance with '
        'the rules of the American Arbitration Association.',
        '8.3 The prevailing party in any arbitration or litigation arising under this '
        'Agreement shall be entitled to recover its reasonable attorneys\' fees and costs.'
    ]
    for text in s8_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Section 9: Miscellaneous ---
    doc.add_heading('Section 9. Miscellaneous', level=2)
    s9_paras = [
        '9.1 Entire Agreement. This Agreement constitutes the entire agreement between the '
        'parties and supersedes all prior negotiations, representations, and agreements, '
        'whether written or oral.',
        '9.2 Amendment. This Agreement may not be amended except by a written instrument '
        'signed by both parties.',
        '9.3 Severability. If any provision of this Agreement is held to be invalid or '
        'unenforceable, the remaining provisions shall continue in full force and effect.'
    ]
    for text in s9_paras:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # --- Signature Block ---
    doc.add_paragraph()  # spacer
    sig = doc.add_paragraph()
    sig.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of the '
                'date first written above.').font.size = Pt(11)

    doc.add_paragraph()
    company_sig = doc.add_paragraph()
    run = company_sig.add_run('MERIDIAN TECHNOLOGIES, INC.')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph('_________________________________')
    name_line = doc.add_paragraph()
    run = name_line.add_run('Jonathan M. Wheeler, Chief Executive Officer')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()
    emp_sig = doc.add_paragraph()
    run = emp_sig.add_run('EMPLOYEE')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph('_________________________________')
    emp_name = doc.add_paragraph()
    run = emp_name.add_run('Alexandra R. Vasquez')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
