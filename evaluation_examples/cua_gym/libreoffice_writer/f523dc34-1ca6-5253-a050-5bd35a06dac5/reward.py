"""
Reward Script: Shift Schedule Template for Customer Service Department
Task ID: writer_hr_085
Domain: libreoffice_writer
Scoring:
  C1 (0.30) - Table exists with correct structure (8 cols, 4 rows)
  C2 (0.25) - Header row has correct day names (Mon-Sun) + "Shift" label
  C3 (0.20) - Three shift rows with correct labels and time ranges
  C4 (0.10) - Data cells accommodate 3 employee names (>=3 lines)
  C5 (0.10) - Week date range header present above the table
  C6 (0.05) - Notes section below table for shift swap requests
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_085'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct structure — 8 cols, 4 rows (0.30 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            # Need 8 columns (Shift + Mon-Sun) and 4 rows (header + 3 shifts)
            if num_cols == 8 and num_rows >= 4:
                print(f"PASS: Component 1 — Table has {num_rows} rows x {num_cols} cols (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected >=4 rows x 8 cols")
        else:
            print(f"FAIL: Component 1 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row has correct day names + "Shift" label (0.25 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            expected_days = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']

            # Check first column says "shift" (or similar)
            has_shift_label = 'shift' in header_cells[0]

            # Check days are present in header (cols 1-7)
            days_found = 0
            for day in expected_days:
                if any(day in cell for cell in header_cells[1:]):
                    days_found += 1

            if has_shift_label and days_found == 7:
                print(f"PASS: Component 2 — Header row has 'Shift' and all 7 day names (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — shift_label={has_shift_label}, days_found={days_found}/7, headers={header_cells}")
        else:
            print(f"FAIL: Component 2 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Shift rows with correct labels and time ranges (0.20 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            shift_checks = 0

            # Expected shifts with time ranges
            expected_shifts = [
                ('morning', '6am', '2pm'),
                ('afternoon', '2pm', '10pm'),
                ('night', '10pm', '6am'),
            ]

            for row_idx in range(1, min(4, len(table.rows))):
                shift_text = table.rows[row_idx].cells[0].text.strip().lower()
                for name, start, end in expected_shifts:
                    # Normalize for matching: remove spaces, colons
                    normalized = shift_text.replace(' ', '').replace(':', '')
                    if name in normalized and start.replace(' ', '') in normalized and end.replace(' ', '') in normalized:
                        shift_checks += 1
                        break

            if shift_checks == 3:
                print(f"PASS: Component 3 — All 3 shifts with correct labels and times (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Only {shift_checks}/3 shifts correctly labeled")
                # Partial credit: give proportional score
                partial = round(0.20 * shift_checks / 3, 2)
                if partial > 0:
                    print(f"  Partial credit: {partial} pts")
                    total_score += partial
        else:
            print(f"FAIL: Component 3 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Data cells accommodate 3 employee names (>=3 lines per cell) (0.10 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            cells_with_3_slots = 0
            total_data_cells = 0

            for row_idx in range(1, min(4, len(table.rows))):
                for col_idx in range(1, min(8, len(table.columns))):
                    total_data_cells += 1
                    cell_text = table.rows[row_idx].cells[col_idx].text.strip()
                    # Count number of paragraphs or newline-separated lines
                    cell_paras = table.rows[row_idx].cells[col_idx].paragraphs
                    num_lines = len([p for p in cell_paras if p.text.strip()])
                    if num_lines >= 3:
                        cells_with_3_slots += 1

            if total_data_cells > 0 and cells_with_3_slots >= total_data_cells * 0.8:
                print(f"PASS: Component 4 — {cells_with_3_slots}/{total_data_cells} data cells have >=3 name slots (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Only {cells_with_3_slots}/{total_data_cells} data cells have >=3 name slots")
        else:
            print(f"FAIL: Component 4 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Week date range header present above the table (0.10 points)
    # In golden, this is a paragraph containing "Week of:" before the table
    try:
        # Find text before the table that references a week/date range
        week_header_matches = [
            p for p in doc.paragraphs
            if re.search(r'week\s*(of|:|\s)', p.text.strip().lower())
            or re.search(r'\b(march|april|may|june|july|august|september|october|november|december|january|february)\b.*\d', p.text.strip().lower())
        ]

        if len(week_header_matches) > 0:
            print(f"PASS: Component 5 — Week date range header found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — No week date range header found in paragraphs")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Notes section below table for shift swap requests (0.05 points)
    try:
        # Look for paragraphs that mention notes/swap
        notes_matches = [
            p for p in doc.paragraphs
            if ('note' in p.text.strip().lower() or 'swap' in p.text.strip().lower())
            and len(p.text.strip()) > 3
        ]

        if len(notes_matches) > 0:
            print(f"PASS: Component 6 — Notes/shift swap section found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No notes or shift swap section found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
