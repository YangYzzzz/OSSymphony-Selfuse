"""
Initial Setup: Weekly Shift Schedule template - pre-task state
Task ID: writer_hr_085
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_085'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: Landscape orientation ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # --- Document Title ---
    title = doc.add_heading('Weekly Shift Schedule', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Department Information ---
    dept_info = doc.add_paragraph()
    dept_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    dept_info.paragraph_format.space_after = Pt(2)
    run = dept_info.add_run('Department: Customer Service')
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.bold = True

    mgr_info = doc.add_paragraph()
    mgr_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    mgr_info.paragraph_format.space_after = Pt(2)
    run = mgr_info.add_run('Manager: Patricia Alvarez')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    loc_info = doc.add_paragraph()
    loc_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    loc_info.paragraph_format.space_after = Pt(2)
    run = loc_info.add_run('Location: Building C, 2nd Floor')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    contact_info.paragraph_format.space_after = Pt(6)
    run = contact_info.add_run('Contact: ext. 4271 | cs-scheduling@company.com')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Separator line ---
    sep = doc.add_paragraph()
    sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sep.add_run('_' * 80)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(8)

    # --- Placeholder instruction ---
    instr = doc.add_paragraph()
    instr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    instr.paragraph_format.space_before = Pt(12)
    run = instr.add_run('[Shift schedule table to be added below]')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
