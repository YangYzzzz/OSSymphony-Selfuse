"""
Initial Setup: Poetry chapbook with 8 poems, no table of contents
Task ID: writer_creative_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_038'
# Context says file is on ~/Desktop/
OUTPUT = f'{WORKDIR}/Desktop/chapbook.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a page break as a separate paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)
    return para


def add_poem(doc, title, body_lines):
    """Add a poem with its title heading and body stanzas."""
    # Poem title as a heading
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_para.paragraph_format.space_after = Pt(12)

    # Poem body stanzas
    for stanza in body_lines:
        if stanza == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph(stanza)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)


def create_initial():
    # Ensure Desktop directory exists on VM
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document()

    # Set page margins (standard 1-inch margins)
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Remove the default empty paragraph if present
    # The document starts with poems - page 1 is the first poem page
    # but context says poems start on page 2 (after a first page which we'll call page 1 with poem 1)
    # Actually context: "Poems start on page 2 of the current document"
    # So page 1 is some intro/first content, poems start at page 2
    # But the task says there are 8 poems each starting on a new page
    # Let's interpret: first poem is on page 1 (no front matter in initial state),
    # poems occupy pages 1-8 (8 poems, page breaks between them).
    # After TOC is inserted, poems shift to pages 2-9.
    # The context says "Poems start on page 2" meaning in initial state page 1 is poem 1.

    # No default empty paragraph to remove in python-docx 1.2.0

    # Poem data: realistic poem content
    poems = [
        (
            "First Snow",
            [
                "The first flakes fall before the city wakes,",
                "soft lanterns drifting past the window glass.",
                "Each rooftop holds its breath and silence makes",
                "a map of what we were, and what must pass.",
                "",
                "By noon the avenues wear white fur coats,",
                "footprints pressed like signatures in slate.",
                "A child cups snow and watches as it floats",
                "through open palms, unable now to wait.",
                "",
                "At dusk the plows arrive with salt and noise,",
                "and all that pristine quiet melts to grey.",
                "But for one morning, winter kept its poise—",
                "and we stood still, and let the world delay.",
            ]
        ),
        (
            "Harbor Lights",
            [
                "The harbor keeps its lanterns burning low,",
                "a string of amber beads along the shore.",
                "The fishing boats rock gently, hull to hull,",
                "while cormorants stand sentry on the pier.",
                "",
                "I used to come here when the fog was thick,",
                "to watch the light dissolve into the sound.",
                "My father said the sea remembers everything—",
                "each name of every sailor that it drowned.",
                "",
                "The tide comes in and rearranges stones.",
                "The harbor light revolves, indifferent, slow.",
                "Some distances remain. The foghorn moans.",
                "I drive back home along the coast road, alone.",
            ]
        ),
        (
            "Grandmother's Garden",
            [
                "She kept the roses ruthlessly in check,",
                "pruning each cane before the first thaw came.",
                "Her hands wore leather gloves with worn-down thumbs,",
                "and every thorn she knew by feel and name.",
                "",
                "The dahlias she staked against the wind,",
                "the peonies she wrapped in paper bags.",
                "She had no patience for disorder—",
                "even weeds she pulled in morning's coolest drags.",
                "",
                "What blooms she left me are not in the ground.",
                "They live in how I handle difficult things,",
                "the way I turn my anger into action,",
                "the care I take with small and tender things.",
            ]
        ),
        (
            "The Last Train",
            [
                "The platform empties in the autumn rain.",
                "One umbrella turned inside out by wind.",
                "The last train out was forty minutes late—",
                "or so the digital sign insistently pinned.",
                "",
                "I met you here in summer, years before,",
                "your luggage stacked like arguments I'd lost.",
                "The commuters parted, the PA crackled,",
                "and we stood counting everything it cost.",
                "",
                "Tonight I watch a stranger check his phone",
                "and wonder what dispatch he's waiting for.",
                "The train arrives. He steps aboard. Moves on.",
                "I remain here, remembering the shore.",
            ]
        ),
        (
            "Wildflowers",
            [
                "Nobody planted these—they simply came,",
                "pushing through the gravel and the clay.",
                "Queen Anne's lace and chicory and yarrow,",
                "asserting their small claim on the highway.",
                "",
                "I think of all the things that grow uninvited:",
                "the volunteer tomato in the crack,",
                "the moss that colonizes shaded stone,",
                "the vine that always finds a way to track.",
                "",
                "Not every living thing needs to be curated.",
                "Some beauty insists without permission first.",
                "The wildflower takes what no one thought to offer—",
                "the margin space, the rocky ground, the worst.",
            ]
        ),
        (
            "Sunday Morning",
            [
                "The newspaper spreads across the kitchen table.",
                "Coffee steams in the winter morning light.",
                "Outside, the birch trees lean their silver shoulders",
                "against the cold, and somewhere sparrows fight",
                "",
                "over the suet cage we hung last Thursday.",
                "My mother calls to say she dreamed of Dad.",
                "I tell her I dreamed nothing, which is easier",
                "than saying what the actual dream I had.",
                "",
                "The coffee cools. The paper stays unread.",
                "The sparrows win or lose—it's hard to say.",
                "I think about the things we don't inherit,",
                "and all the things we carry anyway.",
            ]
        ),
        (
            "Letters from Abroad",
            [
                "She wrote on thin blue paper, the kind that folds",
                "into its own envelope, a secret tongue.",
                "The stamps were bright: a temple, a parade.",
                "Her handwriting looped like something young.",
                "",
                "She wrote of markets loud with morning noise,",
                "of bread and spice and unfamiliar skies.",
                "She wrote of how the light fell differently there—",
                "gold at dusk, and lavender at rise.",
                "",
                "I kept each letter in a shoebox lid.",
                "Years later, moving, I found them still intact.",
                "Her voice came back so clearly, so precisely,",
                "that I sat down on the floor and did not act.",
            ]
        ),
        (
            "Coming Home",
            [
                "The key still fits, though the lock resists a moment.",
                "The hallway smells of paint and someone's past.",
                "The rooms have been rearranged by other hands,",
                "but certain shadows fall exactly as they last.",
                "",
                "I find the mark I made at age of nine",
                "still penciled faint beside the kitchen door.",
                "My height. The date. My mother's name beside it.",
                "The floor still creaks in the same worn-in floor.",
                "",
                "Home is not the rooms but what persists",
                "beneath the surface change: the angle of the light,",
                "the way the garden makes its seasonal return,",
                "the weight of memory that makes a house a site.",
            ]
        ),
    ]

    for i, (title, lines) in enumerate(poems):
        add_poem(doc, title, lines)
        # Add page break after each poem except the last
        if i < len(poems) - 1:
            add_page_break(doc)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open with LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
