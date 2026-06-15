"""
Initial Setup: Create a Writer document with a visible 'Advanced Troubleshooting' section
and a document variable 'UserLevel' set to 'admin'.
Task ID: writer_tech_082
Domain: libreoffice_writer

The document is a technical troubleshooting guide with 7 sections.
Section 6 ('Advanced Troubleshooting') is visible (NOT conditional yet).
A document variable 'UserLevel' = 'admin' is pre-set.
Format: ODT (required for conditional sections, an ODF feature).
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_082'
OUTPUT_DOCX = f'{WORKDIR}/{TASK_ID}.docx'
OUTPUT_ODT = f'{WORKDIR}/{TASK_ID}.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, env=env,
    )
    time.sleep(delay_sec)


def run_cmd(command, timeout=60):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    result = subprocess.run(
        command, shell=True, capture_output=True, text=True,
        timeout=timeout, env=env
    )
    if result.stdout.strip():
        print(f"STDOUT: {result.stdout.strip()}")
    if result.stderr.strip():
        print(f"STDERR: {result.stderr.strip()}")
    return result


def create_base_docx():
    """Create the base .docx document with all content visible."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    doc.add_heading('Network Infrastructure Troubleshooting Guide', level=0)
    meta = doc.add_paragraph()
    meta.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = meta.add_run('Version 3.2 | Last Updated: March 2026 | IT Operations Division')
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph('')

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This guide provides step-by-step instructions for diagnosing and resolving '
        'common network infrastructure issues encountered in our enterprise environment. '
        'It covers Layer 2 switching, Layer 3 routing, DNS resolution, DHCP services, '
        'and VPN connectivity problems.'
    )
    doc.add_paragraph(
        'All procedures should be performed by authorized IT staff. Ensure you have '
        'appropriate access credentials before beginning any troubleshooting session.'
    )

    doc.add_heading('2. Basic Connectivity Checks', level=1)
    doc.add_paragraph('Before diving into complex diagnostics, verify basic connectivity:')
    for item in [
        'Confirm physical cable connections and link lights on switches',
        'Verify IP address assignment via ipconfig /all (Windows) or ip addr (Linux)',
        'Test local gateway reachability: ping 10.0.1.1',
        'Test DNS resolution: nslookup internal.corp.example.com',
        'Test external connectivity: ping 8.8.8.8',
        'Check for packet loss: ping -c 100 10.0.1.1 and review statistics',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3. DHCP Troubleshooting', level=1)
    doc.add_paragraph('If a client fails to obtain an IP address, follow these steps:')
    for i, step in enumerate([
        'Release the current lease: ipconfig /release',
        'Renew the lease: ipconfig /renew',
        'If renewal fails, check the DHCP server status on srv-dhcp-01 (10.0.1.10)',
        'Verify the DHCP scope has available addresses in the 10.0.1.100-10.0.1.250 range',
        'Check for rogue DHCP servers using Wireshark DHCP filter: bootp.type == 2',
        'Review DHCP server logs at /var/log/dhcpd.log for DHCPDECLINE messages',
    ], 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('4. DNS Resolution Issues', level=1)
    doc.add_paragraph(
        'DNS failures are among the most common network issues. Primary DNS: '
        'dns-primary.corp.example.com (10.0.1.5), Secondary: dns-secondary.corp.example.com (10.0.1.6).'
    )
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['Symptom', 'Likely Cause', 'Resolution']):
        cell = table.cell(0, i)
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
    for r_idx, row_data in enumerate([
        ['Cannot resolve internal names', 'DNS server unreachable', 'Verify connectivity to 10.0.1.5'],
        ['Slow DNS resolution', 'DNS cache corruption', 'Flush DNS cache: ipconfig /flushdns'],
        ['Intermittent failures', 'Network congestion', 'Run traceroute to DNS server'],
        ['Wrong IP returned', 'Stale DNS record', 'Clear record; wait for TTL expiry'],
    ], 1):
        for c_idx, val in enumerate(row_data):
            table.cell(r_idx, c_idx).text = val
    doc.add_paragraph('')

    doc.add_heading('5. VPN Connectivity', level=1)
    doc.add_paragraph(
        'Remote users connecting via GlobalProtect VPN may experience connectivity issues. '
        'VPN concentrator: vpn-gw.corp.example.com (203.0.113.50).'
    )
    for item in [
        'Authentication failures: Verify AD credentials and MFA token synchronization',
        'Split-tunnel misconfiguration: Check routing table for 10.0.0.0/8 via VPN tunnel',
        'MTU issues: Set MTU to 1400 on the VPN adapter',
        'Certificate expiration: Renew client certificate from PKI portal',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6. Advanced Troubleshooting', level=1)
    doc.add_paragraph(
        'This section contains advanced diagnostic procedures intended for senior network '
        'engineers and system administrators with elevated access privileges.'
    )
    doc.add_heading('6.1 Packet Capture Analysis', level=2)
    doc.add_paragraph(
        'Use tcpdump or Wireshark to capture traffic on the affected segment. '
        'For targeted captures on the core switch (sw-core-01):'
    )
    doc.add_paragraph('    tcpdump -i eth0 -w /tmp/capture.pcap -c 10000 host 10.0.1.100')
    doc.add_paragraph(
        'Analyze the capture for TCP retransmissions, RST packets, and unusual '
        'traffic patterns. Filter in Wireshark: tcp.analysis.retransmission'
    )
    doc.add_heading('6.2 Spanning Tree Protocol Diagnostics', level=2)
    doc.add_paragraph('STP issues can cause network loops and broadcast storms. On Cisco IOS:')
    for cmd in [
        'show spanning-tree summary  -- overview of STP instances',
        'show spanning-tree interface gi0/1 detail  -- per-port STP state',
        'show spanning-tree blockedports  -- identify blocked ports',
        'debug spanning-tree events  -- real-time STP event logging',
    ]:
        doc.add_paragraph(cmd, style='List Bullet')
    doc.add_heading('6.3 BGP Route Analysis', level=2)
    doc.add_paragraph(
        'For WAN issues, examine BGP peering. Border router: rtr-border-01 (AS 65001).'
    )
    for cmd in [
        'show ip bgp summary  -- verify peer state',
        'show ip bgp neighbors 203.0.113.1 received-routes  -- check prefixes',
        'show ip route bgp  -- verify BGP routes in routing table',
        'clear ip bgp 203.0.113.1 soft in  -- soft reset without dropping session',
    ]:
        doc.add_paragraph(cmd, style='List Bullet')
    doc.add_heading('6.4 Firewall Rule Audit', level=2)
    doc.add_paragraph(
        'When traffic is unexpectedly blocked, review firewall rules on fw-perimeter-01 '
        '(Palo Alto PA-5260). Export the current ruleset and compare against the approved '
        'baseline stored in the CMDB.'
    )
    doc.add_paragraph(
        'Check hit counts on deny rules: show rule-hit-count vsys1 rules. Pay special '
        'attention to rules with implicit deny actions.'
    )

    doc.add_heading('7. Escalation Procedures', level=1)
    doc.add_paragraph('If the issue cannot be resolved, escalate as follows:')
    esc_table = doc.add_table(rows=5, cols=4)
    esc_table.style = 'Table Grid'
    for i, h in enumerate(['Team', 'Email', 'Phone', 'Scope']):
        cell = esc_table.cell(0, i)
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
    for r_idx, row_data in enumerate([
        ['Tier 1 - Help Desk', 'helpdesk@corp.example.com', 'ext. 1000', 'Basic connectivity'],
        ['Tier 2 - Network Ops', 'netops@corp.example.com', 'ext. 2000', 'Switch/router issues'],
        ['Tier 3 - Network Eng', 'neteng@corp.example.com', 'ext. 3000', 'BGP, firewall, architecture'],
        ['Vendor - Cisco TAC', 'N/A', '1-800-553-2447', 'Hardware failures, IOS bugs'],
    ], 1):
        for c_idx, val in enumerate(row_data):
            esc_table.cell(r_idx, c_idx).text = val

    doc.save(OUTPUT_DOCX)
    print(f'Base DOCX created: {OUTPUT_DOCX}')


def convert_and_add_variable():
    """
    Use UNO to:
    1. Open the .docx
    2. Add UserLevel='admin' document variable
    3. Save as ODT (to preserve ODF features for the agent to work with)
    """
    run_cmd("killall soffice.bin 2>/dev/null || true")
    time.sleep(2)

    uno_script = r'''#!/usr/bin/env python3
import subprocess, time, os

env = os.environ.copy()
env["DISPLAY"] = ":0"
subprocess.Popen(
    ["soffice", "--headless", "--invisible", "--norestore",
     "--accept=socket,host=localhost,port=2002;urp;StarOffice.ServiceManager"],
    env=env, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
)
print("Started soffice listener...")
time.sleep(5)

import uno
from com.sun.star.beans import PropertyValue

def connect():
    ctx_local = uno.getComponentContext()
    resolver = ctx_local.ServiceManager.createInstanceWithContext(
        "com.sun.star.bridge.UnoUrlResolver", ctx_local)
    for i in range(10):
        try:
            ctx = resolver.resolve(
                "uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext")
            smgr = ctx.ServiceManager
            return smgr.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
        except Exception as e:
            print(f"  connect attempt {i+1}: {e}")
            time.sleep(2)
    raise RuntimeError("Cannot connect")

try:
    desktop = connect()
    print("Connected")

    prop = PropertyValue()
    prop.Name = "Hidden"
    prop.Value = True
    doc = desktop.loadComponentFromURL(
        "file:///home/user/writer_tech_082.docx", "_blank", 0, (prop,))
    assert doc is not None, "Failed to open document"
    print("Document opened")

    # Set UserLevel variable
    masters = doc.getTextFieldMasters()
    mname = "com.sun.star.text.fieldmaster.User.UserLevel"
    if masters.hasByName(mname):
        master = masters.getByName(mname)
    else:
        master = doc.createInstance("com.sun.star.text.fieldmaster.User")
        master.Name = "UserLevel"
    master.Content = "admin"
    doc.getTextFields().refresh()
    print("UserLevel = admin")

    # Save as ODT (native ODF format preserves all features)
    fp = PropertyValue()
    fp.Name = "FilterName"
    fp.Value = "writer8"
    doc.storeToURL("file:///home/user/writer_tech_082.odt", (fp,))
    doc.close(True)
    print("Saved as ODT")

    with open("/tmp/initial_done.txt", "w") as f:
        f.write("OK")

except Exception as e:
    print(f"ERROR: {e}")
    import traceback; traceback.print_exc()
    with open("/tmp/initial_error.txt", "w") as f:
        f.write(str(e))
finally:
    os.system("killall soffice.bin 2>/dev/null")
    print("Done")
'''

    script_path = '/tmp/convert_initial.py'
    with open(script_path, 'w') as f:
        f.write(uno_script)

    for p in ['/tmp/initial_done.txt', '/tmp/initial_error.txt']:
        if os.path.exists(p):
            os.remove(p)

    print("Converting to ODT and adding UserLevel variable...")
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    result = subprocess.run(
        ['python3', script_path],
        capture_output=True, text=True, timeout=90, env=env
    )
    print(f"STDOUT: {result.stdout}")
    if result.stderr:
        print(f"STDERR: {result.stderr}")

    if os.path.exists('/tmp/initial_done.txt'):
        print("Conversion successful")
    elif os.path.exists('/tmp/initial_error.txt'):
        with open('/tmp/initial_error.txt') as f:
            print(f"Error: {f.read()}")

    run_cmd("killall soffice.bin 2>/dev/null || true")
    time.sleep(2)


def create_initial():
    # Step 1: Create the base docx with python-docx
    create_base_docx()

    # Step 2: Convert to ODT and add UserLevel variable via UNO
    convert_and_add_variable()

    # Step 3: Clean up the intermediate docx
    if os.path.exists(OUTPUT_DOCX):
        os.remove(OUTPUT_DOCX)

    # Step 4: GUI-ready startup
    if os.path.exists(OUTPUT_ODT):
        launch_gui(f'libreoffice --writer "{OUTPUT_ODT}"', delay_sec=3.0)
        print(f'GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
        print(f'Initial file ready: {OUTPUT_ODT} ({os.path.getsize(OUTPUT_ODT)} bytes)')
    else:
        print('ERROR: ODT file not created!')


create_initial()
