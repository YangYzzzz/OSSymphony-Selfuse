"""
Initial Setup: Invoice document with 8 lines, no tabstops, all left-aligned
Task ID: osworld_writer_tabstop_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set document title in properties
    doc.core_properties.title = "Invoice Draft"

    # 8 invoice lines - each has at least 4 descriptive words followed by a dollar amount
    # All plain text, no tabs, no tabstops, all left-aligned
    invoice_lines = [
        "Professional Web Design Services $1200.00",
        "Monthly Server Hosting Fee $89.99",
        "SSL Certificate Renewal Annual $49.00",
        "Search Engine Optimization Package $350.00",
        "Custom Logo Branding Design $275.50",
        "Email Marketing Campaign Setup $199.00",
        "Database Migration Cloud Storage $450.00",
        "Technical Support Maintenance Contract $125.00",
    ]

    for line in invoice_lines:
        para = doc.add_paragraph()
        # No tab stops set - plain left-aligned text with no tabs
        run = para.add_run(line)
        run.font.size = Pt(12)
        # No tabstops, no tabs in text - exactly as described in initial state

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
