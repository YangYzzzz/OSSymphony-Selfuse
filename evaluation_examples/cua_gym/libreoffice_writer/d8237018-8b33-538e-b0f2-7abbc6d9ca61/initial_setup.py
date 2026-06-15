"""
Initial Setup: Apply strikethrough and double-underline formatting for revision markup
Task ID: writer_rd_087
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_087'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Service Agreement — Revision Draft", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle / metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run("Between Apex Consulting Group and Meridian Technologies Inc.")
    run.font.size = Pt(12)
    run.italic = True

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = meta2.add_run("Effective Date: January 15, 2025  |  Revision: March 28, 2025")
    run2.font.size = Pt(10)

    doc.add_paragraph("")  # spacer

    # --- Revision Note ---
    note = doc.add_paragraph()
    rn = note.add_run("REVISION NOTE: ")
    rn.bold = True
    rn.font.size = Pt(11)
    note.add_run(
        "The following document contains the revised text of the Service Agreement. "
        "Changes from the original agreement are annotated below. Phrases that were "
        "removed from the original are noted with [DELETED] comments, and phrases that "
        "were newly added are noted with [INSERTED] comments. Please apply visual "
        "formatting to clearly distinguish deletions and insertions."
    ).font.size = Pt(11)

    doc.add_paragraph("")

    # --- Section 1: Scope of Services ---
    h1 = doc.add_heading("1. Scope of Services", level=2)

    p1 = doc.add_paragraph()
    p1.add_run(
        "Apex Consulting Group (hereinafter \"the Provider\") agrees to deliver "
        "comprehensive strategic consulting services to Meridian Technologies Inc. "
        "(hereinafter \"the Client\"). The engagement shall encompass "
    ).font.size = Pt(11)
    # DELETED phrase #1
    p1.add_run("quarterly performance reviews").font.size = Pt(11)
    p1.add_run(
        " [DELETED — was in original, removed in revision] "
        "and ongoing advisory support for operational optimization. "
        "The Provider shall assign a dedicated team of "
    ).font.size = Pt(11)
    # INSERTED phrase #1
    p1.add_run("senior-level analysts with domain expertise").font.size = Pt(11)
    p1.add_run(
        " [INSERTED — new addition in revision] to ensure continuity and quality."
    ).font.size = Pt(11)

    # --- Section 2: Compensation and Payment Terms ---
    h2 = doc.add_heading("2. Compensation and Payment Terms", level=2)

    p2 = doc.add_paragraph()
    p2.add_run(
        "The Client shall compensate the Provider at a rate of $275 per hour for all "
        "consulting services rendered. Invoices will be submitted "
    ).font.size = Pt(11)
    # DELETED phrase #2
    p2.add_run("on the last business day of each month").font.size = Pt(11)
    p2.add_run(
        " [DELETED — was in original, removed in revision] "
        "and payment shall be due within 30 calendar days of receipt. "
        "A late payment surcharge of "
    ).font.size = Pt(11)
    # DELETED phrase #3
    p2.add_run("2.5% per month").font.size = Pt(11)
    p2.add_run(
        " [DELETED — was in original, removed in revision] "
        "shall apply to outstanding balances. The revised terms now specify that invoices "
        "will be submitted "
    ).font.size = Pt(11)
    # INSERTED phrase #2
    p2.add_run("bi-weekly on the 1st and 15th of each month").font.size = Pt(11)
    p2.add_run(
        " [INSERTED — new addition in revision] with a simplified penalty structure."
    ).font.size = Pt(11)

    # --- Section 3: Confidentiality ---
    h3 = doc.add_heading("3. Confidentiality and Non-Disclosure", level=2)

    p3 = doc.add_paragraph()
    p3.add_run(
        "Both parties agree to maintain strict confidentiality regarding all proprietary "
        "information exchanged during the course of this engagement. The obligation of "
        "confidentiality shall extend for a period of "
    ).font.size = Pt(11)
    # DELETED phrase #4
    p3.add_run("three (3) years").font.size = Pt(11)
    p3.add_run(
        " [DELETED — was in original, removed in revision] "
        "following the termination of this agreement. Confidential information includes, "
        "but is not limited to, financial records, client databases, strategic plans, "
        "and "
    ).font.size = Pt(11)
    # INSERTED phrase #3
    p3.add_run("proprietary algorithms and machine learning models").font.size = Pt(11)
    p3.add_run(
        " [INSERTED — new addition in revision] developed during the engagement."
    ).font.size = Pt(11)

    # --- Section 4: Termination ---
    h4 = doc.add_heading("4. Termination Clause", level=2)

    p4 = doc.add_paragraph()
    p4.add_run(
        "Either party may terminate this agreement by providing "
    ).font.size = Pt(11)
    # DELETED phrase #5
    p4.add_run("sixty (60) days written notice").font.size = Pt(11)
    p4.add_run(
        " [DELETED — was in original, removed in revision] "
        "to the other party. Upon termination, the Provider shall deliver all completed "
        "work products and documentation within fifteen (15) business days. Any "
        "outstanding invoices shall remain payable according to the terms specified in "
        "Section 2 above."
    ).font.size = Pt(11)

    # --- Section 5: Signatures ---
    doc.add_paragraph("")
    h5 = doc.add_heading("5. Governing Law", level=2)
    p5 = doc.add_paragraph()
    p5.add_run(
        "This agreement shall be governed by and construed in accordance with the laws "
        "of the State of California. Any disputes arising under this agreement shall be "
        "resolved through binding arbitration in San Francisco, California."
    ).font.size = Pt(11)

    doc.add_paragraph("")
    sig = doc.add_paragraph()
    sig.add_run("_________________________").font.size = Pt(11)
    sig.add_run("                    ").font.size = Pt(11)
    sig.add_run("_________________________").font.size = Pt(11)

    sig2 = doc.add_paragraph()
    sig2.add_run("Elena Vasquez, Partner").font.size = Pt(10)
    sig2.add_run("                              ").font.size = Pt(10)
    sig2.add_run("David Park, CTO").font.size = Pt(10)

    sig3 = doc.add_paragraph()
    sig3.add_run("Apex Consulting Group").font.size = Pt(10)
    sig3.add_run("                          ").font.size = Pt(10)
    sig3.add_run("Meridian Technologies Inc.").font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
