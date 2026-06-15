"""
Initial Setup: Meeting minutes document with ordered list (no blank lines between items)
Task ID: osworld_writer_blank_line_insertion_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_blank_line_insertion_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Q2 Planning Meeting Minutes', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Meeting Info ---
    doc.add_paragraph('Date: March 12, 2025')
    doc.add_paragraph('Time: 10:00 AM – 11:30 AM')
    doc.add_paragraph('Location: Conference Room B / Zoom')
    doc.add_paragraph('')

    # --- Attendees ---
    doc.add_heading('Attendees', level=1)
    doc.add_paragraph('Sarah Chen – Product Manager (Chair)')
    doc.add_paragraph('Marcus Johnson – Engineering Lead')
    doc.add_paragraph('Priya Sharma – UX Designer')
    doc.add_paragraph('David Kim – QA Engineer')
    doc.add_paragraph('Elena Torres – Marketing Manager')
    doc.add_paragraph('James Okafor – Data Analyst')
    doc.add_paragraph('')

    # --- Agenda ---
    doc.add_heading('Agenda', level=1)
    doc.add_paragraph('Q1 Performance Review', style='List Number')
    doc.add_paragraph('Q2 Product Roadmap Discussion', style='List Number')
    doc.add_paragraph('Resource Allocation and Budget Review', style='List Number')
    doc.add_paragraph('Risk Assessment for Q2 Deliverables', style='List Number')
    doc.add_paragraph('')

    # --- Discussion Notes ---
    doc.add_heading('Discussion Notes', level=1)

    doc.add_heading('1. Q1 Performance Review', level=2)
    doc.add_paragraph(
        'Sarah presented the Q1 dashboard showing a 14% increase in active users '
        'and a 9% improvement in retention compared to Q4 2024. Revenue targets '
        'were met at 97% achievement. Marcus noted three critical bugs were resolved '
        'before the March release, which contributed to the improved NPS score of 42.'
    )

    doc.add_heading('2. Q2 Product Roadmap', level=2)
    doc.add_paragraph(
        'The team reviewed the updated roadmap. Key milestones include the mobile '
        'redesign launch on April 28, the analytics dashboard v2 release on May 15, '
        'and the enterprise SSO integration planned for June 10. Priya shared mockups '
        'for the new onboarding flow, which received positive feedback from the group.'
    )

    doc.add_heading('3. Resource Allocation', level=2)
    doc.add_paragraph(
        'James presented the updated budget projections. Engineering headcount '
        'remains at 12 FTEs, with one contractor backfill approved for the June '
        'integration sprint. Marketing budget for Q2 is set at $85,000, with $30,000 '
        'earmarked for the April product launch campaign.'
    )

    doc.add_heading('4. Risk Assessment', level=2)
    doc.add_paragraph(
        'David highlighted three open risks: (1) dependency on third-party SSO vendor '
        'timeline, (2) potential scope creep in the analytics v2 feature set, and '
        '(3) limited QA capacity in May due to two team members on leave. Mitigation '
        'plans were discussed and assigned.'
    )

    doc.add_paragraph('')

    # --- Action Items (ordered list, NO blank lines between items) ---
    doc.add_heading('Action Items', level=1)

    # 7 list items with NO blank lines between them
    doc.add_paragraph(
        'Marcus to finalize engineering sprint plan for mobile redesign by March 18.',
        style='List Number'
    )
    doc.add_paragraph(
        'Priya to deliver updated onboarding flow prototypes to the team by March 20.',
        style='List Number'
    )
    doc.add_paragraph(
        'Sarah to schedule individual 1:1 sessions with each team lead before March 22.',
        style='List Number'
    )
    doc.add_paragraph(
        'James to prepare Q2 budget summary report and circulate for review by March 19.',
        style='List Number'
    )
    doc.add_paragraph(
        'David to document Q2 risk register and share with stakeholders by March 21.',
        style='List Number'
    )
    doc.add_paragraph(
        'Elena to draft the Q2 marketing campaign brief and present at next meeting.',
        style='List Number'
    )
    doc.add_paragraph(
        'All team leads to confirm resource availability for June SSO integration sprint by March 25.',
        style='List Number'
    )

    doc.add_paragraph('')

    # --- Next Meeting ---
    doc.add_heading('Next Meeting', level=1)
    doc.add_paragraph('Date: March 26, 2025')
    doc.add_paragraph('Time: 10:00 AM – 11:00 AM')
    doc.add_paragraph('Location: Conference Room B / Zoom')
    doc.add_paragraph('Agenda: Sprint review, roadmap updates, and Q2 launch readiness check.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
