"""
Reward Script: Insert merged title row in table
Task ID: writer_tm_042
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table has 9 rows (added 1 row)
  Component 2 (0.30): Row 0 cells are merged spanning 4 columns
  Component 3 (0.20): Row 0 text is 'Financial Summary Q4 2025'
  Component 4 (0.10): Row 0 paragraph alignment is CENTER
  Component 5 (0.15): Row 0 text run is bold
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_042'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def persist_app_state():
    """Try to save any unsaved LibreOffice state via Ctrl+S."""
    try:
        import time
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table has 9 rows (0.25 points)
    # Initial table has 8 rows; the task adds 1 row on top -> 9 rows total
    try:
        if num_rows == 9:
            print(f"PASS: Component 1 — Table has 9 rows (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected 9 rows, found {num_rows}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row 0 is merged across all 4 columns (gridSpan=4) (0.30 points)
    # In initial env, row 0 has no merge (gridSpan=1 per cell).
    # In golden env, row 0 cell 0 has gridSpan=4.
    try:
        row0 = table.rows[0]
        cell0_tc = row0.cells[0]._tc
        grid_span_elems = cell0_tc.findall(f'.//{{{W_NS}}}gridSpan')
        if grid_span_elems:
            span_val = int(grid_span_elems[0].get(f'{{{W_NS}}}val', '1'))
        else:
            span_val = 1

        if span_val >= 4:
            print(f"PASS: Component 2 — Row 0 cell 0 gridSpan={span_val} (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — Row 0 cell 0 gridSpan={span_val}, expected >=4")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Row 0 text is 'Financial Summary Q4 2025' (0.20 points)
    # Initial env row 0 text is 'Category' (the old header). Golden has the title.
    try:
        row0_text = table.rows[0].cells[0].text.strip()
        if row0_text == 'Financial Summary Q4 2025':
            print(f"PASS: Component 3 — Row 0 text matches title (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — Row 0 text is '{row0_text}', expected 'Financial Summary Q4 2025'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Row 0 paragraph alignment is CENTER (0.10 points)
    # Initial env row 0 has no center alignment (headers are default/left).
    try:
        cell0 = table.rows[0].cells[0]
        alignment = cell0.paragraphs[0].paragraph_format.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            print(f"PASS: Component 4 — Row 0 is center-aligned (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Row 0 alignment is {alignment}, expected CENTER")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Row 0 contains 'Financial Summary Q4 2025' AND it is bold (0.15 points)
    # Compound check: anchored to the task-introduced title text being bold.
    # Initial env row 0 has 'Category' (not the title), so this fails on initial.
    try:
        cell0 = table.rows[0].cells[0]
        title_bold = False
        for para in cell0.paragraphs:
            for run in para.runs:
                if 'Financial Summary' in run.text and run.bold:
                    title_bold = True
                    break
            if title_bold:
                break

        if title_bold:
            print(f"PASS: Component 5 — Title 'Financial Summary Q4 2025' is bold (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Title text not found or not bold in row 0")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
