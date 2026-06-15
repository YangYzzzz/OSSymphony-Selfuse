"""
Initial Setup: Insert a title row above a financial data table
Task ID: writer_tm_042
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_042'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a brief intro paragraph for realism
    intro = doc.add_paragraph('Below is the quarterly financial overview for the Operations division.')
    intro.paragraph_format.space_after = Pt(12)

    # Create 4x8 table (4 columns, 8 rows: 1 header + 7 data)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'

    # Row 1: Headers
    headers = ['Category', 'Budget', 'Actual', 'Variance']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(11)

    # Rows 2-8: Financial data (7 data rows)
    data = [
        ['Personnel Costs',    '$245,000', '$238,450', '-$6,550'],
        ['Equipment Leases',   '$87,500',  '$91,200',  '+$3,700'],
        ['Software Licenses',  '$42,000',  '$39,750',  '-$2,250'],
        ['Office Supplies',    '$15,800',  '$16,340',  '+$540'],
        ['Travel & Training',  '$63,200',  '$58,900',  '-$4,300'],
        ['Utilities',          '$28,500',  '$29,100',  '+$600'],
        ['Miscellaneous',      '$12,000',  '$14,260',  '+$2,260'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
