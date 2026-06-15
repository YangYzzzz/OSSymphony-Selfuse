"""
Reward Script: Add footer with 'HR Department - Page' and page number
Task ID: writer_hr_010
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Footer is active (has content, not empty/linked-to-previous default)
  Component 2 (0.4): Footer text contains 'HR Department - Page'
  Component 3 (0.3): Footer contains PAGE field code for auto page numbering
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_010'

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # We check across ALL sections; the footer must appear in at least one section
    footer_active = False
    footer_text_match = False
    page_field_found = False

    for i, section in enumerate(doc.sections):
        footer = section.footer

        # Collect footer text from all paragraphs in this section's footer
        all_footer_text = ""
        has_fld_page = False

        if footer.paragraphs:
            for fp in footer.paragraphs:
                all_footer_text += fp.text

                # Check for PAGE field code in this paragraph
                instr_texts = fp._element.findall('.//w:instrText', NS)
                for it in instr_texts:
                    if it.text and 'PAGE' in it.text.upper():
                        has_fld_page = True

        # Also check footer XML directly for field codes (in case they're outside paragraphs)
        if not has_fld_page:
            instr_texts = footer._element.findall('.//w:instrText', NS)
            for it in instr_texts:
                if it.text and 'PAGE' in it.text.upper():
                    has_fld_page = True

        # Determine if this section has an active footer with content
        if all_footer_text.strip():
            footer_active = True

        # Check text content
        if 'HR Department' in all_footer_text and 'Page' in all_footer_text:
            footer_text_match = True

        if has_fld_page:
            page_field_found = True

    # Component 1: Footer is active with content (0.3 points)
    # This checks that a footer exists and has non-empty text.
    # Initial file has no footer content → FAIL; Golden has footer text → PASS
    try:
        if footer_active:
            print(f"PASS: Component 1 — Footer is active with content (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No active footer with content found in any section")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Footer text contains 'HR Department - Page' (0.4 points)
    # The task requires the footer to show 'HR Department - Page' followed by page number.
    # Initial has empty footer → FAIL; Golden has 'HR Department - Page ' → PASS
    try:
        if footer_text_match:
            print(f"PASS: Component 2 — Footer contains 'HR Department' and 'Page' text (0.4 pts)")
            total_score += 0.4
        else:
            # Collect actual text for debug
            actual_texts = []
            for sec in doc.sections:
                for fp in sec.footer.paragraphs:
                    if fp.text.strip():
                        actual_texts.append(fp.text)
            print(f"FAIL: Component 2 — Expected 'HR Department' and 'Page' in footer, found: {actual_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Footer contains PAGE field code for auto page numbering (0.3 points)
    # The task requires an auto page number field, not just static text.
    # Initial has no field codes → FAIL; Golden has PAGE instrText → PASS
    try:
        if page_field_found:
            print(f"PASS: Component 3 — PAGE field code found in footer (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — No PAGE field code found in any section footer")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
