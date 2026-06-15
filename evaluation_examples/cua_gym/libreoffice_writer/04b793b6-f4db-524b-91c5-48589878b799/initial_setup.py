"""
Initial Setup: Attendance Policy document with 5 pages and no footer
Task ID: writer_hr_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # === PAGE 1 ===
    title = doc.add_heading('Attendance Policy', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Solutions Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: January 1, 2025')
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph()  # spacer

    doc.add_heading('1. Purpose', level=1)
    doc.add_paragraph(
        'This policy establishes the guidelines and expectations for employee attendance '
        'at Meridian Solutions Inc. Regular and punctual attendance is essential for '
        'maintaining productivity, meeting project deadlines, and ensuring effective '
        'collaboration across all departments. This policy applies to all full-time, '
        'part-time, and contract employees.'
    )
    doc.add_paragraph(
        'The company recognizes that employees may occasionally need time away from work '
        'for personal, medical, or family reasons. This policy provides a framework that '
        'balances organizational needs with employee well-being while maintaining fair '
        'and consistent standards for all team members.'
    )

    doc.add_heading('2. Scope', level=1)
    doc.add_paragraph(
        'This policy covers all employees across every office location, including the '
        'headquarters in Portland, Oregon, and regional offices in Austin, Texas; '
        'Denver, Colorado; and Raleigh, North Carolina. Remote employees are also '
        'subject to these attendance expectations during their scheduled working hours.'
    )

    # === PAGE 2 ===
    doc.add_page_break()

    doc.add_heading('3. Standard Work Hours', level=1)
    doc.add_paragraph(
        'Standard business hours are 9:00 AM to 5:30 PM, Monday through Friday. '
        'Employees are expected to be at their workstation or logged in remotely '
        'by their scheduled start time. A 30-minute lunch break is provided between '
        '12:00 PM and 1:00 PM, though departments may adjust this window with '
        'manager approval.'
    )

    doc.add_heading('3.1 Flexible Schedule Arrangements', level=2)
    doc.add_paragraph(
        'Eligible employees may request flexible scheduling through the HR portal. '
        'Approved flex schedules include early start (7:00 AM - 3:30 PM), late start '
        '(10:00 AM - 6:30 PM), or compressed work weeks (four 10-hour days). All '
        'flex schedule requests require written manager approval and must be renewed '
        'quarterly.'
    )

    doc.add_heading('3.2 Remote Work Days', level=2)
    doc.add_paragraph(
        'Employees in eligible roles may work remotely up to two days per week. '
        'Remote days must be pre-approved by the direct supervisor and logged in '
        'the attendance tracking system by the preceding Friday. Remote employees '
        'must remain accessible via Slack and email during standard business hours.'
    )

    doc.add_heading('4. Reporting Absences', level=1)
    doc.add_paragraph(
        'Employees must notify their direct supervisor of any planned absence at '
        'least 48 hours in advance. For unplanned absences due to illness or '
        'emergency, employees must contact their supervisor within one hour of '
        'their scheduled start time. Notification should be made via phone call '
        'or the HR attendance portal; text messages alone are not considered '
        'sufficient notification.'
    )

    # === PAGE 3 ===
    doc.add_page_break()

    doc.add_heading('5. Paid Time Off (PTO)', level=1)
    doc.add_paragraph(
        'Meridian Solutions provides a combined PTO bank that covers vacation, '
        'personal days, and sick leave. Accrual rates are based on length of service:'
    )

    # PTO table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Years of Service', 'Annual PTO Days', 'Monthly Accrual']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    pto_data = [
        ['0 - 2 years', '15 days', '1.25 days'],
        ['3 - 5 years', '20 days', '1.67 days'],
        ['6 - 10 years', '25 days', '2.08 days'],
        ['10+ years', '30 days', '2.50 days'],
    ]
    for r, row_data in enumerate(pto_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacer

    doc.add_heading('5.1 PTO Request Procedures', level=2)
    doc.add_paragraph(
        'All PTO requests must be submitted through the HR portal at least five '
        'business days before the requested dates. Requests for three or more '
        'consecutive days require additional approval from the department director. '
        'During peak business periods (December 15 - January 5 and the last week of '
        'each fiscal quarter), PTO approval is subject to departmental coverage '
        'requirements.'
    )

    doc.add_heading('5.2 Carryover Policy', level=2)
    doc.add_paragraph(
        'Employees may carry over up to five unused PTO days to the following '
        'calendar year. Carryover days must be used by March 31 of the new year. '
        'Unused PTO beyond the carryover limit will be forfeited unless otherwise '
        'required by state law. Employees who leave the company will be paid for '
        'unused accrued PTO up to a maximum of 10 days.'
    )

    # === PAGE 4 ===
    doc.add_page_break()

    doc.add_heading('6. Tardiness', level=1)
    doc.add_paragraph(
        'An employee is considered tardy when arriving more than 5 minutes after '
        'their scheduled start time without prior approval. Repeated tardiness '
        'affects team productivity and may result in progressive disciplinary action:'
    )
    doc.add_paragraph('3 tardiness incidents in a rolling 30-day period: Verbal warning', style='List Bullet')
    doc.add_paragraph('5 incidents in a rolling 60-day period: Written warning', style='List Bullet')
    doc.add_paragraph('7 incidents in a rolling 90-day period: Performance improvement plan', style='List Bullet')
    doc.add_paragraph('Continued pattern after PIP: Possible termination', style='List Bullet')

    doc.add_heading('7. Unexcused Absences', level=1)
    doc.add_paragraph(
        'An absence is considered unexcused when an employee fails to report to work '
        'and does not provide proper notification as outlined in Section 4. Unexcused '
        'absences are treated more seriously than tardiness:'
    )
    doc.add_paragraph('First unexcused absence: Written warning', style='List Bullet')
    doc.add_paragraph('Second unexcused absence within 90 days: Final written warning', style='List Bullet')
    doc.add_paragraph('Third unexcused absence within 180 days: Termination', style='List Bullet')

    doc.add_paragraph(
        'Three consecutive days of unexcused absence without contact will be '
        'considered job abandonment and may result in immediate termination.'
    )

    # === PAGE 5 ===
    doc.add_page_break()

    doc.add_heading('8. Leave of Absence', level=1)
    doc.add_paragraph(
        'Employees may be eligible for extended leave under the Family and Medical '
        'Leave Act (FMLA), which provides up to 12 weeks of unpaid, job-protected '
        'leave per year. Qualifying reasons include the birth or adoption of a child, '
        'care for an immediate family member with a serious health condition, or the '
        "employee's own serious health condition."
    )

    doc.add_heading('8.1 Medical Leave', level=2)
    doc.add_paragraph(
        'Extended medical leave requires documentation from a licensed healthcare '
        'provider. The HR department will work with the employee to determine '
        'appropriate accommodations and return-to-work plans. Medical certifications '
        'must be submitted within 15 calendar days of the leave request.'
    )

    doc.add_heading('9. Accommodations', level=1)
    doc.add_paragraph(
        'Meridian Solutions is committed to providing reasonable accommodations for '
        'employees with disabilities or medical conditions that affect attendance. '
        'Employees requiring accommodation should contact the HR department to '
        'initiate the interactive process. All accommodation requests are handled '
        'confidentially.'
    )

    doc.add_heading('10. Acknowledgment', level=1)
    doc.add_paragraph(
        'By continuing employment at Meridian Solutions Inc., all employees '
        'acknowledge that they have read, understood, and agree to comply with '
        'this attendance policy. This policy is subject to periodic review and '
        'may be updated at the discretion of the Human Resources department.'
    )

    ack_line = doc.add_paragraph()
    ack_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = ack_line.add_run('\nEmployee Signature: _________________________    Date: ___________')
    run.font.size = Pt(11)

    mgr_line = doc.add_paragraph()
    mgr_line.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = mgr_line.add_run('Manager Signature:  _________________________    Date: ___________')
    run.font.size = Pt(11)

    # Ensure NO footer exists (default Document has no footer, but be explicit)
    for sec in doc.sections:
        sec.footer.is_linked_to_previous = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
