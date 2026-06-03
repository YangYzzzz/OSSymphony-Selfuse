"""
Initial Setup: Writer document with a long article paragraph (500 words)
Task ID: writer_fs_011
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_011'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


ARTICLE_TEXT = (
    "The rapid advancement of renewable energy technologies has fundamentally "
    "transformed the global energy landscape over the past two decades. Solar "
    "photovoltaic systems, once considered prohibitively expensive for widespread "
    "adoption, have experienced dramatic cost reductions that now make them "
    "competitive with traditional fossil fuel sources in many regions. Wind "
    "turbine technology has similarly matured, with offshore installations "
    "reaching capacities that were unimaginable just fifteen years ago. These "
    "developments have been driven by a combination of government incentives, "
    "technological innovation, and increasing public awareness of climate change.\n\n"
    "The integration of energy storage solutions has addressed one of the most "
    "significant challenges facing renewable energy: intermittency. Lithium-ion "
    "battery technology, bolstered by advances in materials science and "
    "manufacturing processes, has enabled utilities to store excess energy "
    "generated during peak production hours for use during periods of low "
    "generation. Pumped hydroelectric storage continues to provide large-scale "
    "energy balancing capabilities, while emerging technologies such as "
    "compressed air energy storage and hydrogen fuel cells offer promising "
    "alternatives for long-duration storage needs.\n\n"
    "The economic implications of this energy transition extend far beyond the "
    "power sector itself. New manufacturing facilities for solar panels, wind "
    "turbines, and battery systems have created thousands of jobs in communities "
    "that previously relied on coal mining or oil extraction. Training programs "
    "and workforce development initiatives have helped displaced workers "
    "transition into these growing industries. Meanwhile, declining energy costs "
    "have benefited consumers and businesses alike, with some regions experiencing "
    "electricity prices that are a fraction of what they were a decade ago.\n\n"
    "Environmental monitoring has revealed encouraging trends in regions that "
    "have aggressively pursued renewable energy adoption. Air quality "
    "improvements in cities that have reduced their reliance on coal-fired power "
    "plants have been particularly notable, with measurable decreases in "
    "respiratory illness rates and associated healthcare costs. Water resources "
    "have also benefited, as renewable energy generation requires significantly "
    "less water than thermal power plants, reducing stress on already strained "
    "freshwater supplies in arid regions.\n\n"
    "Looking forward, the convergence of artificial intelligence, advanced "
    "materials research, and distributed energy systems promises to accelerate "
    "the transition even further. Smart grid technologies powered by machine "
    "learning algorithms are optimizing energy distribution in real time, "
    "reducing waste and improving reliability. Perovskite solar cells and "
    "next-generation wind turbine designs are pushing efficiency boundaries "
    "beyond what current silicon-based panels and conventional turbines can "
    "achieve. Community energy cooperatives and microgrids are empowering local "
    "communities to take control of their energy futures, fostering resilience "
    "against extreme weather events and grid disruptions that are becoming "
    "increasingly common in a changing climate."
)


def create_initial():
    doc = Document()

    # Set default page margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Add the article heading
    heading = doc.add_heading("The Future of Renewable Energy", level=1)
    for run in heading.runs:
        run.font.size = Pt(18)

    # Add the long article text as paragraphs
    for paragraph_text in ARTICLE_TEXT.split('\n\n'):
        para = doc.add_paragraph(paragraph_text.strip())
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.name = "Liberation Serif"
            run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
