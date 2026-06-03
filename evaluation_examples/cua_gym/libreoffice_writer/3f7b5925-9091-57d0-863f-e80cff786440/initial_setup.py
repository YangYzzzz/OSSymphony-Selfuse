"""
Initial Setup: Create a 3-page meeting preparation document
Task ID: writer_obj_057
Domain: libreoffice_writer

Creates meeting_prep.docx at /home/user/Desktop/meeting_prep.docx
Page 2 has agenda items but NO callout box/info box.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'meeting_prep'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set up document margins (standard margins)
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ==================== PAGE 1: Cover / Executive Summary ====================
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Q1 2025 Strategic Planning Meeting')
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_para.add_run('Meeting Preparation Packet')
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()

    # Date and location
    details = [
        ('Date:', 'March 18, 2025'),
        ('Time:', '9:00 AM – 5:00 PM'),
        ('Location:', 'Boardroom 3A, Headquarters'),
        ('Chair:', 'Dr. Patricia Nguyen, VP Strategy'),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        run_label = p.add_run(label + ' ')
        run_label.bold = True
        run_label.font.size = Pt(12)
        p.add_run(value).font.size = Pt(12)

    doc.add_paragraph()

    # Executive summary heading
    h = doc.add_paragraph()
    h_run = h.add_run('Executive Summary')
    h_run.bold = True
    h_run.font.size = Pt(14)
    h_run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    exec_text = (
        'This meeting packet has been prepared for the Q1 Strategic Planning session. '
        'All attendees are expected to review this document prior to arrival. '
        'The agenda covers departmental performance reviews, budget allocations for '
        'the upcoming quarter, new product initiatives, and alignment on company-wide goals.'
    )
    p_exec = doc.add_paragraph(exec_text)
    p_exec.paragraph_format.space_after = Pt(8)

    # Attendees
    attendees_h = doc.add_paragraph()
    att_run = attendees_h.add_run('Confirmed Attendees')
    att_run.bold = True
    att_run.font.size = Pt(12)
    att_run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    attendees = [
        'Dr. Patricia Nguyen – VP Strategy (Chair)',
        'James Okafor – CFO',
        'Sophia Martinez – VP Product',
        'David Kim – VP Engineering',
        'Rachel Thompson – VP Sales',
        'Marcus Williams – VP Marketing',
        'Anita Patel – VP Human Resources',
        'Chris Brennan – Chief Data Officer',
    ]
    for att in attendees:
        doc.add_paragraph(att, style='List Bullet')

    # Page break after page 1
    doc.add_page_break()

    # ==================== PAGE 2: Meeting Agenda ====================
    agenda_h = doc.add_paragraph()
    agenda_run = agenda_h.add_run('Meeting Agenda')
    agenda_run.bold = True
    agenda_run.font.size = Pt(16)
    agenda_run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    date_p = doc.add_paragraph()
    date_p.add_run('Date: March 18, 2025  |  Duration: 8 hours  |  Location: Boardroom 3A').font.size = Pt(10)

    doc.add_paragraph()

    agenda_items = [
        ('9:00 – 9:15 AM', 'Opening Remarks & Housekeeping', 'Dr. Patricia Nguyen'),
        ('9:15 – 10:00 AM', 'Q4 2024 Performance Review', 'James Okafor (CFO)'),
        ('10:00 – 10:45 AM', 'Engineering Roadmap: Key Deliverables', 'David Kim'),
        ('10:45 – 11:00 AM', 'Break', '—'),
        ('11:00 – 11:45 AM', 'Product Strategy & New Initiatives', 'Sophia Martinez'),
        ('11:45 AM – 12:30 PM', 'Sales Pipeline & Q1 Targets', 'Rachel Thompson'),
        ('12:30 – 1:30 PM', 'Lunch Break', '—'),
        ('1:30 – 2:15 PM', 'Marketing Campaigns & Brand Positioning', 'Marcus Williams'),
        ('2:15 – 3:00 PM', 'HR: Talent Acquisition & Retention', 'Anita Patel'),
        ('3:00 – 3:45 PM', 'Data Analytics & Business Intelligence', 'Chris Brennan'),
        ('3:45 – 4:00 PM', 'Break', '—'),
        ('4:00 – 4:45 PM', 'Cross-Departmental Alignment & OKRs', 'All VPs'),
        ('4:45 – 5:00 PM', 'Wrap-Up & Action Items', 'Dr. Patricia Nguyen'),
    ]

    # Table for agenda
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['Time', 'Agenda Item', 'Presenter']
    for i, hdr in enumerate(headers):
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(hdr)
        run.bold = True
        run.font.size = Pt(11)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3564')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for time_slot, item, presenter in agenda_items:
        row_cells = table.add_row().cells
        row_cells[0].text = time_slot
        row_cells[1].text = item
        row_cells[2].text = presenter
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # Pre-reading note (plain paragraph, not a text frame)
    prereading_h = doc.add_paragraph()
    pr_run = prereading_h.add_run('Pre-Reading Materials')
    pr_run.bold = True
    pr_run.font.size = Pt(12)
    pr_run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    prereading_items = [
        'Q4 2024 Financial Summary Report (attached)',
        'Engineering Roadmap Document v3.2 (attached)',
        'Sales Pipeline Analysis – February 2025 (attached)',
        'HR Talent Report – Q4 2024 (attached)',
    ]
    for item in prereading_items:
        doc.add_paragraph(item, style='List Bullet')

    # Page break after page 2
    doc.add_page_break()

    # ==================== PAGE 3: Supporting Information ====================
    support_h = doc.add_paragraph()
    support_run = support_h.add_run('Supporting Information & Reference Material')
    support_run.bold = True
    support_run.font.size = Pt(16)
    support_run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    doc.add_paragraph()

    # Key Metrics section
    metrics_h = doc.add_paragraph()
    metrics_run = metrics_h.add_run('Key Performance Metrics – Q4 2024')
    metrics_run.bold = True
    metrics_run.font.size = Pt(13)

    metrics = [
        ('Total Revenue', '$12.4M', '+8.3% vs Q3 2024'),
        ('Operating Margin', '24.7%', '+1.2pp vs Q3 2024'),
        ('Customer Acquisition', '1,847 new accounts', '+15.2% vs Q3 2024'),
        ('Employee Headcount', '312 FTEs', '+23 vs Q3 2024'),
        ('Product Uptime', '99.94%', 'Target: 99.9%'),
        ('NPS Score', '67', 'Industry avg: 42'),
    ]

    metrics_table = doc.add_table(rows=1, cols=3)
    metrics_table.style = 'Table Grid'
    m_hdr = metrics_table.rows[0].cells
    m_headers = ['Metric', 'Value', 'Change']
    for i, mh in enumerate(m_headers):
        para = m_hdr[i].paragraphs[0]
        run = para.add_run(mh)
        run.bold = True
        run.font.size = Pt(11)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc = m_hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '4472C4')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for metric, value, change in metrics:
        row_cells = metrics_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
        row_cells[2].text = change
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # Upcoming deadlines
    deadlines_h = doc.add_paragraph()
    dl_run = deadlines_h.add_run('Upcoming Deadlines & Action Items')
    dl_run.bold = True
    dl_run.font.size = Pt(13)

    deadlines = [
        'March 25, 2025 – Submit Q1 budget revisions to Finance',
        'March 28, 2025 – Engineering sprint planning kickoff',
        'April 1, 2025 – Q1 board presentation preparation deadline',
        'April 5, 2025 – Performance review cycles open for all managers',
        'April 10, 2025 – Marketing campaign approvals due',
    ]
    for dl in deadlines:
        doc.add_paragraph(dl, style='List Number')

    doc.add_paragraph()

    # Contact information
    contact_h = doc.add_paragraph()
    c_run = contact_h.add_run('Meeting Logistics Contact')
    c_run.bold = True
    c_run.font.size = Pt(12)

    contact_info = (
        'For questions regarding meeting logistics, materials, or schedule changes, '
        'please contact the Executive Coordination team: exec-support@company.com | '
        'Ext. 4210. All materials are confidential and intended for internal use only.'
    )
    doc.add_paragraph(contact_info)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
