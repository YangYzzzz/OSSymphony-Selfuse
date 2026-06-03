"""
Initial Setup: Invoice template document with tab-separated line items (no tab stops set)
Task ID: writer_para_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_para_020'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: 'INVOICE' — Heading 1, center-aligned
    heading = doc.add_heading('INVOICE', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 2: Invoice number
    doc.add_paragraph('Invoice #: INV-2025-0342')

    # Paragraph 3: Date
    doc.add_paragraph('Date: March 1, 2025')

    # Paragraphs 4-8: Line items with tab-separated columns (NO tab stops — agent must add them)
    # Paragraph 4: header row
    doc.add_paragraph('Item\tDescription\tAmount')

    # Paragraph 5: line item 001
    doc.add_paragraph('001\tWeb Development Services - February 2025\t$4,500.00')

    # Paragraph 6: line item 002
    doc.add_paragraph('002\tUI/UX Design Consultation\t$1,200.00')

    # Paragraph 7: line item 003
    doc.add_paragraph('003\tServer Maintenance and Monitoring\t$800.00')

    # Paragraph 8: total row
    doc.add_paragraph('\tTotal:\t$6,500.00')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
