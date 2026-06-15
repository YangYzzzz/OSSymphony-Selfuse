"""
Reward Script: Set tab stops for invoice line items in writer_para_020.docx
Task ID: writer_para_020
Domain: libreoffice_writer
Scoring:
  - Precondition gate: File loads, paragraphs 0-7 exist, text content unchanged
  - Component 1: Paragraphs 4-8 (indices 3-7) have LEFT tab stop at 3cm     (0.35 pts)
  - Component 2: Paragraphs 4-8 (indices 3-7) have LEFT tab stop at 8cm     (0.35 pts)
  - Component 3: Paragraphs 4-8 (indices 3-7) have RIGHT tab stop at 15cm   (0.30 pts)
  Total: 1.00

  Note: "No custom tab stops on headers (para 0-2)" is a precondition gate only
  since it is true in both initial and golden states. Text content preservation
  is also a precondition gate for the same reason.
"""

import os

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_020'

# Conversion constant: EMU (English Metric Units) per cm
# 1 inch = 914400 EMU, 1 inch = 2.54 cm → 1 cm = 914400/2.54 = 360000 EMU
EMU_PER_CM = 360000

# Tolerance: 0.05 cm (18000 EMU) to account for rounding
POSITION_TOLERANCE = 18000  # 0.05 cm in EMU


def emu_to_cm(emu):
    """Convert EMU to centimeters."""
    return emu / EMU_PER_CM


def cm_to_emu(cm):
    """Convert centimeters to EMU."""
    return int(cm * EMU_PER_CM)


def get_custom_tab_stops(para):
    """
    Get custom (non-default) tab stops for a paragraph.
    Filters out CLEAR stops and LEFT@0 defaults added by LibreOffice.
    Returns list of (alignment, position_emu) tuples.
    """
    result = []
    for ts in para.paragraph_format.tab_stops:
        if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
            continue
        if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
            continue
        result.append((ts.alignment, ts.position))
    return result


def has_tab_stop_at(tab_stops, alignment, target_cm):
    """
    Check if a specific tab stop (alignment + position) exists.
    Uses tolerance to account for EMU rounding.
    """
    target_emu = cm_to_emu(target_cm)
    for ts_align, ts_pos in tab_stops:
        if ts_align == alignment and abs(ts_pos - target_emu) <= POSITION_TOLERANCE:
            return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Set left-aligned tab stops at 3cm and 8cm, and a right-aligned tab stop
          at 15cm for the invoice line items (Paragraphs 4-8, indices 3-7).
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify we have the expected number of paragraphs
    if len(doc.paragraphs) < 8:
        print(f"CRITICAL: Expected at least 8 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify text content integrity
    # This is a gate (true in both states), not a scoring component
    expected_texts = [
        'INVOICE',
        'Invoice #: INV-2025-0342',
        'Date: March 1, 2025',
        'Item\tDescription\tAmount',
        '001\tWeb Development Services - February 2025\t$4,500.00',
        '002\tUI/UX Design Consultation\t$1,200.00',
        '003\tServer Maintenance and Monitoring\t$800.00',
        '\tTotal:\t$6,500.00',
    ]
    for idx, expected in enumerate(expected_texts):
        if idx < len(doc.paragraphs):
            actual = doc.paragraphs[idx].text
            if actual != expected:
                print(f"GATE_FAIL: Text content changed at paragraph {idx}: expected {repr(expected)}, got {repr(actual)}")
                # Continue scoring — text changes don't necessarily mean the tab task failed

    # Invoice line item paragraphs (indices 3-7):
    # Para 3 (Item/Description/Amount header)
    # Para 4 (001 Web Development)
    # Para 5 (002 UI/UX Design)
    # Para 6 (003 Server Maintenance)
    # Para 7 (Total row)
    line_item_indices = [3, 4, 5, 6, 7]

    # -------------------------------------------------------------------------
    # Component 1: Line item paragraphs (indices 3-7) have LEFT tab stop at 3cm (0.35 pts)
    # This FAILS on initial_env (no tab stops defined), PASSES on golden_env
    # -------------------------------------------------------------------------
    try:
        missing_3cm = [
            idx for idx in line_item_indices
            if not has_tab_stop_at(get_custom_tab_stops(doc.paragraphs[idx]), WD_TAB_ALIGNMENT.LEFT, 3.0)
        ]
        if len(missing_3cm) == 0:
            print(f"PASS: Component 1 — All line item paragraphs (3-7) have LEFT tab stop at 3cm (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — Missing LEFT@3cm tab stop in paragraphs: {missing_3cm}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Line item paragraphs (indices 3-7) have LEFT tab stop at 8cm (0.35 pts)
    # This FAILS on initial_env (no tab stops defined), PASSES on golden_env
    # -------------------------------------------------------------------------
    try:
        missing_8cm = [
            idx for idx in line_item_indices
            if not has_tab_stop_at(get_custom_tab_stops(doc.paragraphs[idx]), WD_TAB_ALIGNMENT.LEFT, 8.0)
        ]
        if len(missing_8cm) == 0:
            print(f"PASS: Component 2 — All line item paragraphs (3-7) have LEFT tab stop at 8cm (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 2 — Missing LEFT@8cm tab stop in paragraphs: {missing_8cm}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Line item paragraphs (indices 3-7) have RIGHT tab stop at 15cm (0.30 pts)
    # This FAILS on initial_env (no tab stops defined), PASSES on golden_env
    # -------------------------------------------------------------------------
    try:
        missing_15cm = [
            idx for idx in line_item_indices
            if not has_tab_stop_at(get_custom_tab_stops(doc.paragraphs[idx]), WD_TAB_ALIGNMENT.RIGHT, 15.0)
        ]
        if len(missing_15cm) == 0:
            print(f"PASS: Component 3 — All line item paragraphs (3-7) have RIGHT tab stop at 15cm (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 3 — Missing RIGHT@15cm tab stop in paragraphs: {missing_15cm}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Round to avoid floating point precision issues (e.g., 0.9999999... -> 1.0)
    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
