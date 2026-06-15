"""
Reward Script: Configure abstract paragraph with box border, padding, indent, and alignment
Task ID: writer_para_060
Domain: libreoffice_writer
Scoring:
  - Component 1: Box border present with correct style/color/width on all 4 sides (0.40 pts)
  - Component 2: Border padding = 0.3cm (approx 170 TWIPs) on all 4 sides (0.20 pts)
  - Component 3: Left indent = 1cm and right indent = 1cm on abstract paragraph (0.20 pts)
  - Component 4: Alignment = JUSTIFY on abstract paragraph (0.10 pts)
  - Component 5: No unintended side effects on other paragraphs (gated on abstract having borders) (0.10 pts)
Total: 1.0
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_para_060'

# Unit conversion constants
# 0.75pt border: OOXML w:sz uses 1/8 pt units, so 0.75pt = 6 eighths
EXPECTED_BORDER_SZ = 6
EXPECTED_BORDER_COLOR = '000000'
EXPECTED_BORDER_VAL = 'single'

# 0.3cm padding: 0.3cm in TWIPs = 0.3 * (72/2.54) * 20 ≈ 170 TWIPs; allow ±10
EXPECTED_SPACE_TWIPS = 170
SPACE_TOLERANCE = 10

# 1cm indent: in EMU = 360000; in TWIPs = 567; allow ±10000 EMU or ±20 TWIPs
EXPECTED_IND_EMU = 360000
IND_EMU_TOLERANCE = 10000
EXPECTED_IND_TWIPS = 567
IND_TOLERANCE = 20

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_border_attribs(pBdr_elem, side):
    """Return dict of border attributes for given side (top, left, bottom, right), or None if missing."""
    ns = W_NS
    side_elem = pBdr_elem.find(f'{{{ns}}}' + side)
    if side_elem is None:
        return None
    return {
        'val': side_elem.get(f'{{{ns}}}val'),
        'sz': side_elem.get(f'{{{ns}}}sz'),
        'space': side_elem.get(f'{{{ns}}}space'),
        'color': side_elem.get(f'{{{ns}}}color'),
    }


def count_valid_border_sides(pBdr):
    """
    Count sides where border val=single, sz=6, color=000000.
    Returns (valid_count, issues_list) based on actual API checks.
    """
    sides = ['top', 'left', 'bottom', 'right']
    valid_sides = 0
    issues = []
    for side in sides:
        attribs = get_border_attribs(pBdr, side)
        if attribs is None:
            issues.append(f"  {side}: missing border element")
            continue

        side_issues = []

        # Check style (val)
        if attribs['val'] != EXPECTED_BORDER_VAL:
            side_issues.append(f"  {side} val: expected '{EXPECTED_BORDER_VAL}', got '{attribs['val']}'")

        # Check width (sz = 6 means 0.75pt in 1/8 pt units)
        sz_raw = attribs['sz']
        sz_val = int(sz_raw) if sz_raw is not None else -1
        if sz_val != EXPECTED_BORDER_SZ:
            side_issues.append(f"  {side} sz: expected {EXPECTED_BORDER_SZ} (0.75pt), got {sz_val}")

        # Check color (black = 000000)
        color_val = (attribs['color'] or '').upper().lstrip('#')
        if color_val != EXPECTED_BORDER_COLOR:
            side_issues.append(f"  {side} color: expected '{EXPECTED_BORDER_COLOR}', got '{attribs['color']}'")

        if len(side_issues) == 0:
            valid_sides += 1
        else:
            issues.extend(side_issues)

    return valid_sides, issues


def count_correct_padding_sides(pBdr):
    """
    Count sides where border space ≈ EXPECTED_SPACE_TWIPS (0.3cm).
    Returns (valid_count, issues_list).
    """
    sides = ['top', 'left', 'bottom', 'right']
    valid_sides = 0
    issues = []
    for side in sides:
        attribs = get_border_attribs(pBdr, side)
        if attribs is None:
            issues.append(f"  {side}: missing border element")
            continue

        space_raw = attribs['space']
        space_val = int(space_raw) if space_raw is not None else -1
        if abs(space_val - EXPECTED_SPACE_TWIPS) <= SPACE_TOLERANCE:
            valid_sides += 1
        else:
            issues.append(
                f"  {side} space: expected ~{EXPECTED_SPACE_TWIPS} TWIPs (0.3cm), got {space_val}"
            )
    return valid_sides, issues


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Configure the abstract paragraph (paragraph starting with "We present a novel") with:
    - Box border: 0.75pt solid black on all 4 sides
    - Border padding: 0.3cm on all sides
    - Left indent: 1cm
    - Right indent: 1cm
    - Alignment: JUSTIFY
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Identify abstract paragraph by text prefix (as specified in task context)
    abstract_para = None
    abstract_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("We present a novel deep reinforcement learning framework"):
            abstract_para = para
            abstract_idx = i
            break

    if abstract_para is None:
        print("CRITICAL: Could not locate abstract paragraph in document.")
        print("REWARD: 0.0")
        return 0.0

    print(f"Found abstract paragraph at index {abstract_idx}: '{abstract_para.text[:60]}...'")

    # Pre-fetch XML elements used across components
    pPr = abstract_para._p.find(qn('w:pPr'))
    pBdr = pPr.find(qn('w:pBdr')) if pPr is not None else None

    # -------------------------------------------------------------------------
    # Component 1: Box border present with correct style, color, width on all 4
    # sides (0.40 pts).
    # Verification: val=single, sz=6 (0.75pt in 1/8-pt OOXML units), color=000000
    # This FAILS on initial (no pBdr element) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        if pBdr is None:
            print("FAIL: Component 1 — No paragraph border (pBdr) element found on abstract paragraph")
        else:
            valid_sides, issues = count_valid_border_sides(pBdr)
            if valid_sides == 4:
                print("PASS: Component 1 — Box border on all 4 sides: val=single, sz=6 (0.75pt), color=000000 (0.40 pts)")
                total_score += 0.40
            else:
                print(f"FAIL: Component 1 — Only {valid_sides}/4 sides have correct border. Issues:")
                for issue in issues:
                    print(issue)
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Border padding = 0.3cm (~170 TWIPs) on all 4 sides (0.20 pts)
    # Checks w:space attribute on each border side element.
    # This FAILS on initial (no pBdr element) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        if pBdr is None:
            print("FAIL: Component 2 — No pBdr element, cannot check padding")
        else:
            valid_sides, issues = count_correct_padding_sides(pBdr)
            if valid_sides == 4:
                print(f"PASS: Component 2 — Border padding ≈{EXPECTED_SPACE_TWIPS} TWIPs (0.3cm) on all 4 sides (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — Only {valid_sides}/4 sides have correct padding. Issues:")
                for issue in issues:
                    print(issue)
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Left indent = 1cm and right indent = 1cm (0.20 pts)
    # Checks via python-docx EMU (pf.left_indent) and XML TWIPs as fallback.
    # This FAILS on initial (no indent set) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        pf = abstract_para.paragraph_format
        li = pf.left_indent   # EMU or None
        ri = pf.right_indent  # EMU or None

        # XML-level fallback via w:ind element
        ind_elem = pPr.find(qn('w:ind')) if pPr is not None else None
        xml_left = None
        xml_right = None
        if ind_elem is not None:
            left_attr = ind_elem.get(qn('w:left'))
            right_attr = ind_elem.get(qn('w:right'))
            xml_left = int(left_attr) if left_attr is not None else None
            xml_right = int(right_attr) if right_attr is not None else None

        # Check EMU path first, then TWIPs fallback
        emu_valid = (
            li is not None and ri is not None
            and abs(li - EXPECTED_IND_EMU) <= IND_EMU_TOLERANCE
            and abs(ri - EXPECTED_IND_EMU) <= IND_EMU_TOLERANCE
        )
        twips_valid = (
            xml_left is not None and xml_right is not None
            and abs(xml_left - EXPECTED_IND_TWIPS) <= IND_TOLERANCE
            and abs(xml_right - EXPECTED_IND_TWIPS) <= IND_TOLERANCE
        )

        if emu_valid or twips_valid:
            total_score += 0.20
            print(
                f"PASS: Component 3 — Left indent={li} EMU ({xml_left} TWIPs) and "
                f"right indent={ri} EMU ({xml_right} TWIPs) ≈ 1cm (0.20 pts)"
            )
        else:
            print(
                f"FAIL: Component 3 — Indent values out of range. "
                f"Got left={li} EMU ({xml_left} TWIPs), right={ri} EMU ({xml_right} TWIPs); "
                f"expected ~{EXPECTED_IND_EMU} EMU (~{EXPECTED_IND_TWIPS} TWIPs = 1cm)"
            )
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Alignment = JUSTIFY (0.10 pts)
    # This FAILS on initial (LEFT alignment) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        alignment = abstract_para.paragraph_format.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            print("PASS: Component 4 — Alignment is JUSTIFY (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Expected JUSTIFY ({WD_PARAGRAPH_ALIGNMENT.JUSTIFY}), found: {alignment}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: No unintended side effects on other paragraphs (0.10 pts)
    # Compound check: gated on abstract paragraph having a border (pBdr != None),
    # ensuring this component also FAILS on initial_env where pBdr is absent.
    # When gated: verifies no other paragraph gained an unexpected border or indent.
    # -------------------------------------------------------------------------
    try:
        if pBdr is None:
            # Gate: abstract has no border → initial_env → component fails
            print("FAIL: Component 5 — Gate failed (no border on abstract para); cannot verify side effects")
        else:
            # Count other paragraphs with unexpected borders/large indents
            violation_count = sum(
                1
                for i, para in enumerate(doc.paragraphs)
                if i != abstract_idx
                for pPr_other in [para._p.find(qn('w:pPr'))]
                if pPr_other is not None
                for _ in [None]  # iterate once
                if (
                    pPr_other.find(qn('w:pBdr')) is not None
                    or any(
                        ind_other is not None
                        and (
                            (ind_other.get(qn('w:left')) and int(ind_other.get(qn('w:left'))) > 200)
                            or (ind_other.get(qn('w:right')) and int(ind_other.get(qn('w:right'))) > 200)
                        )
                        for ind_other in [pPr_other.find(qn('w:ind'))]
                    )
                )
            )

            if violation_count == 0:
                print("PASS: Component 5 — No other paragraphs have unexpected borders or indents (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 5 — {violation_count} other paragraph(s) have unexpected borders/indents")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
