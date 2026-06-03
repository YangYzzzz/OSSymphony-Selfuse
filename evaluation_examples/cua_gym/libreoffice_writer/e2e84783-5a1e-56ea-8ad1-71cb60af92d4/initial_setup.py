"""
Initial Setup: Conference paper docx for paragraph formatting task
Task ID: writer_para_060
Domain: libreoffice_writer

Creates /home/user/writer_para_060.docx with the conference paper content.
The abstract paragraph (paragraph 4) has NO box border, NO padding, NO special indent,
and NO justified alignment — those are the task's goal for the agent to apply.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_060'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title — Heading 1, center-aligned
    title_para = doc.add_heading(
        'Deep Reinforcement Learning for Autonomous Navigation in Dynamic Environments',
        level=1
    )
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 2: Authors line — normal paragraph, no special formatting
    authors_para = doc.add_paragraph(
        'Authors: Zhang Wei, Kumar Patel, Sofia Rodriguez \u2014 Department of Robotics, ETH Zurich'
    )

    # Paragraph 3: "Abstract" heading — Heading 2
    doc.add_heading('Abstract', level=2)

    # Paragraph 4: Abstract body — plain paragraph, NO box border, NO padding,
    # NO left/right indent, NO justified alignment (agent will apply these)
    abstract_para = doc.add_paragraph(
        'We present a novel deep reinforcement learning framework for autonomous robot '
        'navigation that adapts in real-time to dynamic obstacles. Our approach combines '
        'model-predictive control with a learned value function, achieving a 94% success '
        'rate in cluttered environments compared to 71% for baseline methods. Extensive '
        'experiments on both simulated and real-world platforms demonstrate the robustness '
        'and generalization capability of our approach across diverse scenarios.'
    )
    # Explicitly set alignment to LEFT (not JUSTIFY) to ensure no pre-applied justify
    abstract_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Paragraph 5: Keywords
    doc.add_paragraph(
        'Keywords: reinforcement learning, autonomous navigation, robotics, obstacle avoidance'
    )

    # Paragraph 6: Section heading
    doc.add_paragraph('1. Introduction')

    # Paragraph 7: Introduction body
    doc.add_paragraph(
        'Autonomous navigation in environments with moving obstacles remains one of the '
        'fundamental challenges in robotics.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
