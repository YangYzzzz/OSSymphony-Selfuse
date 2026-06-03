"""
Initial Setup: Format envelope addresses with specific fonts
Task ID: writer_lec_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_040'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up an envelope-style page (landscape, smaller page)
    section = doc.sections[0]
    section.page_width = Inches(9.5)
    section.page_height = Inches(4.125)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Return address (top-left, smaller) - 12pt Times New Roman
    return_lines = [
        "Greenfield Technologies Inc.",
        "4200 Lakeshore Boulevard, Suite 300",
        "Chicago, IL 60613"
    ]
    for line_text in return_lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Add some blank paragraphs to separate return and delivery address
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)

    # Delivery address (center, larger) - 12pt Times New Roman
    delivery_lines = [
        "Ms. Patricia Hawthorne",
        "Regional Director, Western Operations",
        "Cascade Financial Group",
        "8750 Wilshire Boulevard, Floor 22",
        "Los Angeles, CA 90036"
    ]
    for line_text in delivery_lines:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent = Inches(2.5)
        run = para.add_run(line_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
