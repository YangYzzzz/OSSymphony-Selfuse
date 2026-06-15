"""
Initial Setup: Insert continuous section break after introduction paragraph
Task ID: writer_page_050
Domain: libreoffice_writer

Creates a 5-page popular science article (science_article.docx) on the Desktop.
The document has a single-column layout throughout (no section breaks).
The agent must insert a continuous section break after the introduction paragraph
and apply a 2-column layout to the body text section.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_050'
OUTPUT = f'{WORKDIR}/science_article.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup: A4, portrait, all margins 2.54cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("The Invisible Engine: How Microbes Power Life on Earth")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Times New Roman"
    title_para.paragraph_format.space_after = Pt(12)
    title_para.paragraph_format.space_before = Pt(0)

    # Byline
    byline_para = doc.add_paragraph()
    byline_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    byline_run = byline_para.add_run("Dr. Elena Vasquez, Department of Microbiology, Stanford University")
    byline_run.italic = True
    byline_run.font.size = Pt(11)
    byline_run.font.name = "Times New Roman"
    byline_para.paragraph_format.space_after = Pt(18)

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run("Published: March 2025 | Popular Science Digest")
    date_run.font.size = Pt(10)
    date_run.font.name = "Times New Roman"
    date_para.paragraph_format.space_after = Pt(24)

    # INTRODUCTION paragraph (this is the one after which the section break will be inserted)
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(
        "Every breath you take, every meal you digest, and every drop of rain that falls on fertile soil "
        "is touched by the invisible hand of microbial life. Bacteria, archaea, fungi, and viruses—organisms "
        "too small to see with the naked eye—collectively form the most influential ecological force on our "
        "planet. They outnumber the stars in the observable universe, and their combined biomass dwarfs that "
        "of every plant and animal combined. Yet for most of human history, we were blissfully unaware of "
        "their existence. Today, as our tools grow sharper and our understanding deepens, scientists are "
        "uncovering just how profoundly these microscopic architects shape the chemistry of our oceans, the "
        "fertility of our soils, the health of our bodies, and even the composition of the atmosphere we "
        "breathe. This article explores the remarkable world of microbial ecology and the cascading ways "
        "these invisible organisms sustain all complex life on Earth."
    )
    intro_run.font.size = Pt(12)
    intro_run.font.name = "Times New Roman"
    intro_para.paragraph_format.first_line_indent = Cm(0)
    intro_para.paragraph_format.space_after = Pt(12)
    intro_para.paragraph_format.line_spacing = 1.5

    # BODY SECTION — Section 1: The Microbial Census
    heading1 = doc.add_paragraph()
    h1_run = heading1.add_run("The Microbial Census: Numbers Beyond Comprehension")
    h1_run.bold = True
    h1_run.font.size = Pt(14)
    h1_run.font.name = "Times New Roman"
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)

    body1a = doc.add_paragraph()
    b1a_run = body1a.add_run(
        "The sheer scale of microbial life defies easy comprehension. In a single gram of fertile garden "
        "soil, microbiologists estimate there are anywhere from 100 million to one billion individual "
        "bacterial cells, representing tens of thousands of distinct species. If you were to count one "
        "bacterium per second, it would take over three years just to tally the microbes in a single "
        "teaspoon of healthy soil. Scale that up to the entire surface of Earth's landmasses, and the "
        "numbers become astronomical."
    )
    b1a_run.font.size = Pt(12)
    b1a_run.font.name = "Times New Roman"
    body1a.paragraph_format.first_line_indent = Cm(0.63)
    body1a.paragraph_format.space_after = Pt(6)
    body1a.paragraph_format.line_spacing = 1.5

    body1b = doc.add_paragraph()
    b1b_run = body1b.add_run(
        "The human body itself is a microbial metropolis. For decades, scientists cited a ratio of ten "
        "microbial cells to every human cell, but recent recalibrations suggest the ratio is closer to "
        "one-to-one—approximately 38 trillion microbial cells alongside 30 trillion human cells. The "
        "majority of these reside in the gut, where the microbiome performs tasks ranging from vitamin "
        "synthesis and immune regulation to mood modulation via the gut-brain axis. The complexity of "
        "these interactions rivals that of any organ system in the human body."
    )
    b1b_run.font.size = Pt(12)
    b1b_run.font.name = "Times New Roman"
    body1b.paragraph_format.first_line_indent = Cm(0.63)
    body1b.paragraph_format.space_after = Pt(6)
    body1b.paragraph_format.line_spacing = 1.5

    body1c = doc.add_paragraph()
    b1c_run = body1c.add_run(
        "Marine environments harbor their own staggering microbial populations. The world's oceans contain "
        "an estimated 10^29 microbial cells—more than the number of grains of sand on all Earth's beaches. "
        "Marine cyanobacteria, particularly Prochlorococcus, are responsible for producing roughly 20 percent "
        "of the oxygen in our atmosphere. A single species, invisible to the naked eye, sustains roughly one "
        "in five breaths taken by every living organism on Earth."
    )
    b1c_run.font.size = Pt(12)
    b1c_run.font.name = "Times New Roman"
    body1c.paragraph_format.first_line_indent = Cm(0.63)
    body1c.paragraph_format.space_after = Pt(6)
    body1c.paragraph_format.line_spacing = 1.5

    # Section 2: Nutrient Cycling
    heading2 = doc.add_paragraph()
    h2_run = heading2.add_run("Architects of the Nitrogen Cycle")
    h2_run.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.name = "Times New Roman"
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)

    body2a = doc.add_paragraph()
    b2a_run = body2a.add_run(
        "Nitrogen is essential for building proteins, DNA, and all the molecular machinery of life. "
        "Yet atmospheric nitrogen (N2) is chemically inert—unable to be used directly by plants, animals, "
        "or fungi. The bridge between the abundant nitrogen in our atmosphere and the biologically "
        "available forms that sustain life is built almost exclusively by microorganisms. Nitrogen-fixing "
        "bacteria such as Rhizobium, which form intimate partnerships with legume roots, and free-living "
        "species like Azotobacter convert atmospheric nitrogen into ammonia through a metabolically costly "
        "process that requires significant energy."
    )
    b2a_run.font.size = Pt(12)
    b2a_run.font.name = "Times New Roman"
    body2a.paragraph_format.first_line_indent = Cm(0.63)
    body2a.paragraph_format.space_after = Pt(6)
    body2a.paragraph_format.line_spacing = 1.5

    body2b = doc.add_paragraph()
    b2b_run = body2b.add_run(
        "Once fixed, nitrogen cascades through ecosystems via a series of microbially mediated "
        "transformations. Nitrifying bacteria oxidize ammonia to nitrite and then nitrate, forms "
        "readily absorbed by plant roots. Denitrifying bacteria complete the cycle by converting "
        "nitrates back to atmospheric nitrogen gas, preventing the accumulation of reactive nitrogen "
        "compounds that can acidify soils and waterways. This elegant biogeochemical ballet, choreographed "
        "entirely by microbes, has maintained the chemistry of life on Earth for over three billion years."
    )
    b2b_run.font.size = Pt(12)
    b2b_run.font.name = "Times New Roman"
    body2b.paragraph_format.first_line_indent = Cm(0.63)
    body2b.paragraph_format.space_after = Pt(6)
    body2b.paragraph_format.line_spacing = 1.5

    body2c = doc.add_paragraph()
    b2c_run = body2c.add_run(
        "The disruption of microbial nitrogen cycling is among the most pressing environmental challenges "
        "of our era. Industrial nitrogen fixation through the Haber-Bosch process produces synthetic "
        "fertilizers at a rate that now rivals natural biological fixation globally. This flood of "
        "reactive nitrogen into agricultural systems has fed billions but also triggered widespread "
        "eutrophication of rivers, lakes, and coastal zones, creating oxygen-depleted dead zones where "
        "fish and other aquatic life cannot survive. The microbes that once maintained balance are "
        "overwhelmed by the scale of human intervention."
    )
    b2c_run.font.size = Pt(12)
    b2c_run.font.name = "Times New Roman"
    body2c.paragraph_format.first_line_indent = Cm(0.63)
    body2c.paragraph_format.space_after = Pt(6)
    body2c.paragraph_format.line_spacing = 1.5

    # Section 3: Carbon Cycle
    heading3 = doc.add_paragraph()
    h3_run = heading3.add_run("The Microbial Carbon Pump")
    h3_run.bold = True
    h3_run.font.size = Pt(14)
    h3_run.font.name = "Times New Roman"
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)

    body3a = doc.add_paragraph()
    b3a_run = body3a.add_run(
        "The global carbon cycle—the planetary mechanism by which carbon moves between the atmosphere, "
        "oceans, soil, and living organisms—is fundamentally a microbial enterprise. Photosynthetic "
        "microorganisms, particularly phytoplankton in the ocean and cyanobacteria in freshwater, fix "
        "carbon dioxide from the atmosphere at a rate comparable to all terrestrial plants combined. "
        "When these organisms die, their carbon-rich bodies sink toward the ocean floor, sequestering "
        "carbon in sediments where it may remain for millions of years—a process oceanographers call the "
        "biological carbon pump."
    )
    b3a_run.font.size = Pt(12)
    b3a_run.font.name = "Times New Roman"
    body3a.paragraph_format.first_line_indent = Cm(0.63)
    body3a.paragraph_format.space_after = Pt(6)
    body3a.paragraph_format.line_spacing = 1.5

    body3b = doc.add_paragraph()
    b3b_run = body3b.add_run(
        "Meanwhile, soil microbes play the equally critical role of decomposers. Fungi and bacteria "
        "break down dead organic matter—fallen leaves, dead animals, discarded food—into simpler "
        "compounds that can be absorbed by plant roots or released as carbon dioxide and methane back "
        "into the atmosphere. This decomposition process releases the nutrients locked in organic "
        "matter, making them available for the next generation of plants. Without microbial decomposers, "
        "the world's forests would be buried under centuries of accumulated dead wood and leaf litter, "
        "and the nutrients they contain would be locked away permanently."
    )
    b3b_run.font.size = Pt(12)
    b3b_run.font.name = "Times New Roman"
    body3b.paragraph_format.first_line_indent = Cm(0.63)
    body3b.paragraph_format.space_after = Pt(6)
    body3b.paragraph_format.line_spacing = 1.5

    body3c = doc.add_paragraph()
    b3c_run = body3c.add_run(
        "Climate change poses a profound threat to these microbially mediated carbon cycles. As "
        "permafrost thaws in the Arctic and subarctic regions, previously frozen organic matter becomes "
        "accessible to microbial decomposers for the first time in thousands of years. The resulting "
        "release of carbon dioxide and methane—a greenhouse gas 80 times more potent than CO2 over a "
        "20-year timescale—creates a positive feedback loop that could accelerate warming far beyond "
        "current projections. Understanding and modeling this 'microbial climate feedback' has become "
        "one of the most urgent challenges in climate science."
    )
    b3c_run.font.size = Pt(12)
    b3c_run.font.name = "Times New Roman"
    body3c.paragraph_format.first_line_indent = Cm(0.63)
    body3c.paragraph_format.space_after = Pt(6)
    body3c.paragraph_format.line_spacing = 1.5

    # Section 4: Human Health
    heading4 = doc.add_paragraph()
    h4_run = heading4.add_run("Microbes and Human Health: A Complex Partnership")
    h4_run.bold = True
    h4_run.font.size = Pt(14)
    h4_run.font.name = "Times New Roman"
    heading4.paragraph_format.space_before = Pt(12)
    heading4.paragraph_format.space_after = Pt(6)

    body4a = doc.add_paragraph()
    b4a_run = body4a.add_run(
        "The relationship between microbes and human health extends far beyond the familiar narrative of "
        "pathogens causing disease. The human microbiome—the community of microorganisms that colonize "
        "our bodies from birth—is increasingly recognized as a critical determinant of health across "
        "virtually every organ system. The gut microbiome alone produces hundreds of bioactive compounds, "
        "including short-chain fatty acids that fuel intestinal cells, neurotransmitter precursors that "
        "influence brain function, and signaling molecules that calibrate the immune system's response "
        "to threats."
    )
    b4a_run.font.size = Pt(12)
    b4a_run.font.name = "Times New Roman"
    body4a.paragraph_format.first_line_indent = Cm(0.63)
    body4a.paragraph_format.space_after = Pt(6)
    body4a.paragraph_format.line_spacing = 1.5

    body4b = doc.add_paragraph()
    b4b_run = body4b.add_run(
        "Research over the past decade has linked disruptions to the gut microbiome—termed dysbiosis—to "
        "a striking range of conditions including obesity, type 2 diabetes, Crohn's disease, depression, "
        "autism spectrum disorder, and even certain cancers. While causality remains difficult to "
        "establish in many cases, the associations are compelling and have sparked enormous interest in "
        "microbiome-based therapies. Fecal microbiota transplantation (FMT), in which the gut microbiome "
        "of a healthy donor is transferred to a patient, has proven remarkably effective for recurrent "
        "Clostridioides difficile infections and is being investigated for conditions ranging from "
        "inflammatory bowel disease to Parkinson's disease."
    )
    b4b_run.font.size = Pt(12)
    b4b_run.font.name = "Times New Roman"
    body4b.paragraph_format.first_line_indent = Cm(0.63)
    body4b.paragraph_format.space_after = Pt(6)
    body4b.paragraph_format.line_spacing = 1.5

    # Section 5: Conclusion
    heading5 = doc.add_paragraph()
    h5_run = heading5.add_run("Looking Forward: The Microbial Horizon")
    h5_run.bold = True
    h5_run.font.size = Pt(14)
    h5_run.font.name = "Times New Roman"
    heading5.paragraph_format.space_before = Pt(12)
    heading5.paragraph_format.space_after = Pt(6)

    body5a = doc.add_paragraph()
    b5a_run = body5a.add_run(
        "We stand at a pivotal moment in our understanding of microbial life. Technologies such as "
        "metagenomics—the sequencing of all DNA in an environmental sample—have revealed that the vast "
        "majority of microbial species on Earth have never been cultured in a laboratory and remain "
        "essentially unknown to science. The 'dark matter' of the microbial world almost certainly "
        "harbors organisms with metabolic capabilities and ecological roles we have not yet imagined. "
        "Future discoveries in this domain may yield new antibiotics, novel biofuels, innovative "
        "approaches to carbon sequestration, and transformative treatments for chronic diseases."
    )
    b5a_run.font.size = Pt(12)
    b5a_run.font.name = "Times New Roman"
    body5a.paragraph_format.first_line_indent = Cm(0.63)
    body5a.paragraph_format.space_after = Pt(6)
    body5a.paragraph_format.line_spacing = 1.5

    body5b = doc.add_paragraph()
    b5b_run = body5b.add_run(
        "The great ecologist E.O. Wilson once wrote that microorganisms are the foundation upon which "
        "all higher life depends. Every ecosystem, every body of water, every handful of soil is a "
        "microbial civilization whose workings we are only beginning to decipher. As we face the "
        "interconnected crises of climate change, antibiotic resistance, and ecosystem collapse, our "
        "ability to understand and work with—rather than against—the microbial world may prove to be "
        "our most important scientific frontier. The invisible engine of life on Earth has been running "
        "for four billion years; it is past time we learned to read its manual."
    )
    b5b_run.font.size = Pt(12)
    b5b_run.font.name = "Times New Roman"
    body5b.paragraph_format.first_line_indent = Cm(0.63)
    body5b.paragraph_format.space_after = Pt(6)
    body5b.paragraph_format.line_spacing = 1.5

    # References section
    ref_heading = doc.add_paragraph()
    rh_run = ref_heading.add_run("References")
    rh_run.bold = True
    rh_run.font.size = Pt(13)
    rh_run.font.name = "Times New Roman"
    ref_heading.paragraph_format.space_before = Pt(12)
    ref_heading.paragraph_format.space_after = Pt(6)

    references = [
        "Sender, R., Fuchs, S., & Milo, R. (2016). Revised estimates for the number of human and bacteria cells in the body. Cell, 164(3), 337-340.",
        "Falkowski, P. G., Fenchel, T., & Delong, E. F. (2008). The microbial engines that drive Earth's biogeochemical cycles. Science, 320(5879), 1034-1039.",
        "Macy, J. M., & Schröder, I. (2018). The nitrogen cycle and its impact on human health and environmental sustainability. Annual Review of Microbiology, 72, 147-169.",
        "Turnbaugh, P. J., et al. (2006). An obesity-associated gut microbiome with increased capacity for energy harvest. Nature, 444(7122), 1027-1031.",
        "Schuur, E. A. G., et al. (2015). Climate change and the permafrost carbon feedback. Nature, 520(7546), 171-179.",
    ]
    for ref in references:
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run(ref)
        ref_run.font.size = Pt(10)
        ref_run.font.name = "Times New Roman"
        ref_para.paragraph_format.first_line_indent = Cm(-0.5)
        ref_para.paragraph_format.left_indent = Cm(0.5)
        ref_para.paragraph_format.space_after = Pt(4)

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
