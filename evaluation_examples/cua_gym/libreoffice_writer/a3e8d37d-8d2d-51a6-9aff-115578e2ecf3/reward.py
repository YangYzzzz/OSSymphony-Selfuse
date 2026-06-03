"""
Reward Script: Insert a continuous section break after the introduction paragraph
               so that two-column layout applies only to the body text.
Task ID: writer_page_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.40): A continuous section break is inserted after paragraph 3
                       (the introduction paragraph). Verified by presence of sectPr
                       in para 3's pPr with start_type == CONTINUOUS (0).
  Component 2 (0.30): The body section (section after the break) has a 2-column
                       layout with equalWidth enabled.
  Component 3 (0.30): The column spacing is approximately 0.50cm (180000 EMU ± 5%)
                       and the introduction section is single column.
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_050'
FILE_PATH = f'{WORKDIR}/science_article.docx'

# Tolerance for spacing check: ±5% of 180000 EMU
SPACING_TARGET_EMU = 180000
SPACING_TOLERANCE = 0.05  # 5%


def persist_app_state():
    """Send Ctrl+S to save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that a continuous section break has been inserted after the introduction
    paragraph, and that the body section uses a 2-column layout.

    Returns a float between 0.0 and 1.0.
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have paragraphs
    if not doc.paragraphs:
        print("CRITICAL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Continuous section break inserted after the introduction
    #              paragraph (paragraph index 3 — "Every breath you take...")
    #              (0.40 points)
    # -----------------------------------------------------------------------
    # A section break is represented by a sectPr element inside a paragraph's
    # pPr. When present on paragraph N, it defines the end of that section.
    # A continuous break has start_type == 0 (CONTINUOUS).
    # -----------------------------------------------------------------------
    try:
        intro_para_index = 3  # "Every breath you take..." — introduction paragraph

        # Verify the paragraph at index 3 is indeed the introduction
        if len(doc.paragraphs) <= intro_para_index:
            print(f"FAIL: Component 1 — Document has fewer than {intro_para_index + 1} paragraphs")
        else:
            intro_para = doc.paragraphs[intro_para_index]
            intro_text_snippet = intro_para.text[:30] if intro_para.text else ""

            # Check if paragraph has an inline sectPr (section break marker)
            pPr = intro_para._p.find(qn('w:pPr'))
            inline_sectPr = None
            if pPr is not None:
                inline_sectPr = pPr.find(qn('w:sectPr'))

            if inline_sectPr is not None:
                # Get the start type of this inline section
                type_elem = inline_sectPr.find(qn('w:type'))
                if type_elem is not None:
                    section_type_val = type_elem.get(qn('w:val'), '')
                else:
                    # No explicit type element means default which may vary;
                    # python-docx represents absence as CONTINUOUS for inline breaks
                    section_type_val = 'continuous'

                section_type_lower = section_type_val.lower()
                if section_type_lower == 'continuous' or section_type_lower == '':
                    print(f"PASS: Component 1 — Continuous section break found after intro para "
                          f"('{intro_text_snippet}...'), type='{section_type_val}' (0.40 pts)")
                    total_score += 0.40
                else:
                    print(f"FAIL: Component 1 — Section break found after intro para but type is "
                          f"'{section_type_val}' (expected 'continuous')")
            else:
                print(f"FAIL: Component 1 — No section break (sectPr) found after intro para "
                      f"('{intro_text_snippet}...')")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: The body section (the section AFTER the break, i.e. the
    #              final document section) has 2-column layout with equalWidth.
    #              (0.30 points)
    # -----------------------------------------------------------------------
    # In python-docx, doc.sections returns all sections. When there are 2
    # sections, sections[0] is the introduction section and sections[1] (or
    # the last section) is the body section.
    # -----------------------------------------------------------------------
    try:
        sections = doc.sections
        if len(sections) < 2:
            print(f"FAIL: Component 2 — Expected 2 sections (intro + body), found {len(sections)}")
        else:
            # The body section is the LAST section in the document
            body_section = sections[-1]
            cols_elem = body_section._sectPr.find(qn('w:cols'))

            if cols_elem is None:
                print("FAIL: Component 2 — Body section has no cols element (single column)")
            else:
                num_cols_str = cols_elem.get(qn('w:num'))
                equal_width_str = cols_elem.get(qn('w:equalWidth'), '1')

                num_cols = int(num_cols_str) if num_cols_str else 1
                equal_width = equal_width_str not in ('0', 'false', 'False')

                if num_cols == 2 and equal_width:
                    print(f"PASS: Component 2 — Body section has 2-column layout with equalWidth "
                          f"(num={num_cols}, equalWidth={equal_width_str}) (0.30 pts)")
                    total_score += 0.30
                elif num_cols == 2:
                    print(f"FAIL: Component 2 — Body section has 2 columns but equalWidth="
                          f"{equal_width_str} (expected equalWidth=1/true)")
                else:
                    print(f"FAIL: Component 2 — Body section has {num_cols} column(s), expected 2")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Column spacing ≈ 0.50cm (180000 EMU ± 5%) AND the
    #              introduction section (sections[0]) is single column.
    #              (0.30 points — awarded as a compound check)
    # -----------------------------------------------------------------------
    try:
        sections = doc.sections
        if len(sections) < 2:
            print(f"FAIL: Component 3 — Expected 2 sections, found {len(sections)}")
        else:
            intro_section = sections[0]
            body_section = sections[-1]

            # Check intro section is single column (no cols element, or num=1)
            intro_cols_elem = intro_section._sectPr.find(qn('w:cols'))
            intro_num_cols = 1  # default: single column
            if intro_cols_elem is not None:
                intro_num_str = intro_cols_elem.get(qn('w:num'))
                if intro_num_str and int(intro_num_str) > 1:
                    intro_num_cols = int(intro_num_str)

            # Check column spacing of body section
            body_cols_elem = body_section._sectPr.find(qn('w:cols'))
            spacing_val = None
            if body_cols_elem is not None:
                space_str = body_cols_elem.get(qn('w:space'))
                if space_str:
                    spacing_val = int(space_str)

            lower_bound = SPACING_TARGET_EMU * (1 - SPACING_TOLERANCE)
            upper_bound = SPACING_TARGET_EMU * (1 + SPACING_TOLERANCE)
            spacing_in_range = (
                spacing_val is not None
                and lower_bound <= spacing_val <= upper_bound
            )

            if intro_num_cols == 1 and spacing_in_range:
                total_score += 0.30
                spacing_cm = spacing_val / 360000  # 1 cm = 360000 EMU (914400 / 2.54)
                print(f"PASS: Component 3 — Intro section is single column AND body column spacing "
                      f"is {spacing_val} EMU (~{spacing_cm:.2f}cm, target 0.50cm) (0.30 pts)")
            elif intro_num_cols != 1:
                print(f"FAIL: Component 3 — Introduction section is NOT single column "
                      f"(cols num={intro_num_cols})")
            elif not spacing_in_range:
                if spacing_val is not None:
                    spacing_cm = spacing_val / 360000  # 1 cm = 360000 EMU
                    print(f"FAIL: Component 3 — Body column spacing is {spacing_val} EMU "
                          f"(~{spacing_cm:.2f}cm), expected ~0.50cm (180000 EMU ± 5%)")
                else:
                    print("FAIL: Component 3 — Body section has no spacing attribute on cols element")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Final score
    # -----------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.1f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state()

if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
