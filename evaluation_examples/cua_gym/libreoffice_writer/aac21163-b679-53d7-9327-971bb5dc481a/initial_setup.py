"""
Initial Setup: Apply strikethrough to the penultimate paragraph of a technical specification
Task ID: osworld_writer_strikethrough_last_para_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Introduction / Overview
    p1 = doc.add_paragraph()
    run1 = p1.add_run(
        "This document describes the technical specification for the NextGen Platform v3.0 "
        "data ingestion subsystem. The subsystem is responsible for receiving, validating, "
        "and routing structured data payloads from upstream provider APIs into the internal "
        "processing pipeline. All interfaces comply with the OpenAPI 3.1 schema and use "
        "TLS 1.3 for transport security."
    )
    run1.font.size = Pt(12)

    # Paragraph 2: Architecture and Components
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "The ingestion subsystem consists of three primary components: the API Gateway, "
        "the Validation Engine, and the Routing Controller. The API Gateway exposes RESTful "
        "endpoints on port 8443 and enforces OAuth 2.0 bearer token authentication. "
        "The Validation Engine performs schema validation against registered JSON Schema "
        "definitions (Draft-07). The Routing Controller applies configurable routing rules "
        "to direct payloads to one of five downstream consumers based on payload type "
        "and priority class."
    )
    run2.font.size = Pt(12)

    # Paragraph 3: Performance Requirements
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        "Performance targets for the production environment are defined as follows: "
        "sustained throughput of 12,000 requests per second with a p99 latency not "
        "exceeding 45 milliseconds under nominal load. The system must maintain 99.95% "
        "uptime as measured over any rolling 30-day window. Horizontal scaling is "
        "supported via Kubernetes HPA with a minimum of 3 replicas and a maximum of "
        "24 replicas. Memory allocation per pod is capped at 2 GiB."
    )
    run3.font.size = Pt(12)

    # Paragraph 4 (PENULTIMATE): Outdated Specifications — NO strikethrough in initial
    p4 = doc.add_paragraph()
    run4 = p4.add_run(
        "Prior to the v2.8 release, the ingestion subsystem relied on a synchronous "
        "blocking I/O model using Apache Tomcat 9.0 with a thread-pool size of 200. "
        "Payload sizes were limited to 512 KB, and the system supported a maximum of "
        "1,500 requests per second. Retry logic was handled client-side with a fixed "
        "backoff interval of 5 seconds. These constraints necessitated manual scaling "
        "operations and resulted in occasional message loss during traffic spikes "
        "exceeding 1,800 requests per second."
    )
    run4.font.size = Pt(12)
    # NOTE: NO strikethrough here — that is the task the agent must perform

    # Paragraph 5: Future Work / Conclusion
    p5 = doc.add_paragraph()
    run5 = p5.add_run(
        "Planned enhancements for the v3.1 release include integration with the "
        "enterprise event streaming platform (Apache Kafka 3.6), support for binary "
        "Protocol Buffers payloads in addition to JSON, and a self-service portal "
        "for dynamic routing rule management. Security hardening tasks include migration "
        "to FIPS 140-3 validated cryptographic modules and implementation of mutual "
        "TLS (mTLS) for all internal service-to-service communication."
    )
    run5.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
