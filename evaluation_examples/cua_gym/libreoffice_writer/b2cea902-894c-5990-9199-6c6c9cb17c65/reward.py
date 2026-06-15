"""
Reward Script: Grandma Rose's Kitchen — Family Recipes Cookbook
Task ID: writer_creative_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Title page — "Grandma Rose's Kitchen" 28pt bold centered
                       + "Family Recipes" 18pt italic centered
  Component 2 (0.25): Table of Contents — "Contents" heading + 4 recipe names
                       with page numbers (tab-separated)
  Component 3 (0.25): Recipe name headings — all 4 recipes 18pt bold centered
  Component 4 (0.25): Page breaks >= 5 AND "Ingredients"/"Directions" subheadings 13pt bold
"""

import os

# python-docx for .docx verification
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_050'
FILE_PATH = '%s/Desktop/grandma_recipes.docx' % WORKDIR


def count_page_breaks(doc):
    """Count manual page break elements in all paragraphs."""
    ns_uri = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for br in run.element.findall('.//{%s}br' % ns_uri):
                btype = br.attrib.get('{%s}type' % ns_uri)
                if btype == 'page':
                    count += 1
    return count


def get_para_size_pt(para):
    """Return font size in points from the first run that has a size set."""
    for run in para.runs:
        if run.font.size is not None:
            return run.font.size.pt
    return None


def para_is_bold(para):
    """Return True if any run in para has bold=True explicitly set."""
    return any(run.font.bold is True for run in para.runs)


def para_is_italic(para):
    """Return True if any run in para has italic=True explicitly set."""
    return any(run.font.italic is True for run in para.runs)


def para_is_centered(para):
    """Return True if paragraph alignment is CENTER."""
    return para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER


def find_paragraphs_exact(doc, search_text):
    """Return list of paragraphs whose stripped text exactly equals search_text."""
    return [para for para in doc.paragraphs if para.text.strip() == search_text]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print('CRITICAL: Cannot load file %s: %s' % (file_path, e))
        print('REWARD: 0.0')
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Title page formatting — 0.25 points
    #   "Grandma Rose's Kitchen": 28pt, bold=True, centered
    #   "Family Recipes":         18pt, italic=True, centered
    # -------------------------------------------------------------------------
    try:
        title_paras = find_paragraphs_exact(doc, "Grandma Rose's Kitchen")
        subtitle_paras = find_paragraphs_exact(doc, 'Family Recipes')

        title_checks_pass = (
            len(title_paras) > 0
            and get_para_size_pt(title_paras[0]) is not None
            and abs(get_para_size_pt(title_paras[0]) - 28.0) < 1.0
            and para_is_bold(title_paras[0])
            and para_is_centered(title_paras[0])
        )

        subtitle_checks_pass = (
            len(subtitle_paras) > 0
            and get_para_size_pt(subtitle_paras[0]) is not None
            and abs(get_para_size_pt(subtitle_paras[0]) - 18.0) < 1.0
            and para_is_italic(subtitle_paras[0])
            and para_is_centered(subtitle_paras[0])
        )

        if title_checks_pass:
            print('PASS: Title "Grandma Rose\'s Kitchen" — 28pt, bold, centered')
        else:
            title_size = get_para_size_pt(title_paras[0]) if title_paras else None
            title_bold = para_is_bold(title_paras[0]) if title_paras else None
            title_center = para_is_centered(title_paras[0]) if title_paras else None
            print('FAIL: Title check — found=%s, size=%s, bold=%s, center=%s' % (
                bool(title_paras), title_size, title_bold, title_center))

        if subtitle_checks_pass:
            print('PASS: Subtitle "Family Recipes" — 18pt, italic, centered')
        else:
            sub_size = get_para_size_pt(subtitle_paras[0]) if subtitle_paras else None
            sub_italic = para_is_italic(subtitle_paras[0]) if subtitle_paras else None
            sub_center = para_is_centered(subtitle_paras[0]) if subtitle_paras else None
            print('FAIL: Subtitle check — found=%s, size=%s, italic=%s, center=%s' % (
                bool(subtitle_paras), sub_size, sub_italic, sub_center))

        if title_checks_pass and subtitle_checks_pass:
            total_score += 0.25
            print('COMPONENT 1 PASS: Title page content correct (+0.25)')
        else:
            print('COMPONENT 1 FAIL: Title page not fully correct')

    except Exception as e:
        print('ERROR: Component 1 — %s' % e)

    # -------------------------------------------------------------------------
    # Component 2: Table of Contents — 0.25 points
    #   "Contents" heading exists
    #   All 4 recipe names appear as TOC entries: "RecipeName\tPageNum"
    # -------------------------------------------------------------------------
    try:
        contents_paras = find_paragraphs_exact(doc, 'Contents')
        toc_heading_present = len(contents_paras) > 0

        if toc_heading_present:
            print('PASS: "Contents" heading found')
        else:
            print('FAIL: "Contents" heading not found')

        recipe_names = ['Chocolate Chip Cookies', 'Apple Pie', 'Banana Bread', 'Lemon Bars']
        toc_entries_found = 0
        for recipe in recipe_names:
            # TOC entry: "Recipe Name\t<digit>" — tab + page number
            matching = [
                para for para in doc.paragraphs
                if recipe in para.text and '\t' in para.text
                and para.text.strip().split('\t')[-1].strip().isdigit()
            ]
            if matching:
                toc_entries_found += 1
                print('PASS: TOC entry found for "%s"' % recipe)
            else:
                print('FAIL: TOC entry with page number not found for "%s"' % recipe)

        toc_entries_all_found = (toc_entries_found == 4)

        if toc_heading_present and toc_entries_all_found:
            total_score += 0.25
            print('COMPONENT 2 PASS: Table of Contents correct (+0.25)')
        else:
            print('COMPONENT 2 FAIL: TOC incomplete (%d/4 entries found)' % toc_entries_found)

    except Exception as e:
        print('ERROR: Component 2 — %s' % e)

    # -------------------------------------------------------------------------
    # Component 3: Recipe heading formatting — 0.25 points
    #   All 4 recipe name headings: 18pt, bold=True, centered
    # -------------------------------------------------------------------------
    try:
        recipe_names = ['Chocolate Chip Cookies', 'Apple Pie', 'Banana Bread', 'Lemon Bars']
        recipe_heading_correct_count = 0

        for recipe in recipe_names:
            # Find paragraphs that are exactly the recipe name (not TOC entries with tabs)
            paras = find_paragraphs_exact(doc, recipe)

            # Check if any of these paragraphs passes all formatting requirements
            recipe_formatted_correctly = any(
                get_para_size_pt(para) is not None
                and abs(get_para_size_pt(para) - 18.0) < 1.0
                and para_is_bold(para)
                and para_is_centered(para)
                for para in paras
            )

            if recipe_formatted_correctly:
                recipe_heading_correct_count += 1
                print('PASS: Recipe heading "%s" — 18pt, bold, centered' % recipe)
            else:
                size_info = get_para_size_pt(paras[0]) if paras else None
                bold_info = para_is_bold(paras[0]) if paras else None
                center_info = para_is_centered(paras[0]) if paras else None
                print('FAIL: Recipe heading "%s" — found=%s, size=%s, bold=%s, center=%s' % (
                    recipe, bool(paras), size_info, bold_info, center_info))

        if recipe_heading_correct_count == 4:
            total_score += 0.25
            print('COMPONENT 3 PASS: All 4 recipe headings correctly formatted (+0.25)')
        else:
            print('COMPONENT 3 FAIL: Only %d/4 recipe headings correctly formatted' % recipe_heading_correct_count)

    except Exception as e:
        print('ERROR: Component 3 — %s' % e)

    # -------------------------------------------------------------------------
    # Component 4: Page breaks (>=5) AND subheading formatting — 0.25 points
    #   At least 5 manual page breaks in the document
    #   All "Ingredients" and "Directions" paragraphs: 13pt, bold=True
    # -------------------------------------------------------------------------
    try:
        pb_count = count_page_breaks(doc)
        print('INFO: Found %d page break(s) in document' % pb_count)

        page_breaks_sufficient = pb_count >= 5
        if page_breaks_sufficient:
            print('PASS: %d page breaks found (>= 5 required)' % pb_count)
        else:
            print('FAIL: Only %d page breaks found (need >= 5)' % pb_count)

        # Check all "Ingredients" and "Directions" subheadings
        subheading_issues = []
        for para in doc.paragraphs:
            if para.text.strip() in ('Ingredients', 'Directions'):
                size_pt = get_para_size_pt(para)
                is_bold = para_is_bold(para)
                if not (size_pt is not None and abs(size_pt - 13.0) < 1.0 and is_bold):
                    subheading_issues.append(
                        '"%s" at wrong format: size=%s, bold=%s' % (para.text.strip(), size_pt, is_bold)
                    )

        subheadings_correct = (len(subheading_issues) == 0)
        if subheadings_correct:
            print('PASS: All "Ingredients"/"Directions" subheadings are 13pt bold')
        else:
            print('FAIL: Subheading format issues: %s' % subheading_issues)

        if page_breaks_sufficient and subheadings_correct:
            total_score += 0.25
            print('COMPONENT 4 PASS: Page breaks and subheading formatting correct (+0.25)')
        else:
            print('COMPONENT 4 FAIL: Page breaks or subheading formatting incorrect')

    except Exception as e:
        print('ERROR: Component 4 — %s' % e)

    # -------------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print('\nScore: %.2f/1.0' % total_score)
    print('REWARD: %.1f' % final_score)
    return final_score


if not os.path.exists(FILE_PATH):
    print('File not found: %s' % FILE_PATH)
    print('REWARD: 0.0')
else:
    verify_task(FILE_PATH)
