"""
Initial Setup: Grandma Rose's Cookbook - Initial State (4 recipes, no formatting/structure)
Task ID: writer_creative_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'writer_creative_050'
OUTPUT = f'{DESKTOP}/grandma_recipes.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Remove default paragraph spacing to keep it natural
    # All text: 12pt Times New Roman, no page breaks, continuous

    recipes = [
        {
            "name": "Chocolate Chip Cookies",
            "ingredients": [
                "2 1/4 cups all-purpose flour",
                "1 tsp baking soda",
                "1 tsp salt",
                "1 cup (2 sticks) butter, softened",
                "3/4 cup granulated sugar",
                "3/4 cup packed brown sugar",
                "2 large eggs",
                "2 tsp vanilla extract",
                "2 cups chocolate chips",
            ],
            "directions": [
                "Preheat oven to 375°F. Line baking sheets with parchment paper.",
                "In a small bowl, combine flour, baking soda and salt. Set aside.",
                "In a large mixer bowl, beat butter and both sugars until creamy. Add eggs one at a time, beating well after each addition. Beat in vanilla.",
                "Gradually blend in the flour mixture. Stir in chocolate chips.",
                "Drop rounded tablespoons of dough onto the prepared baking sheets, spacing them about 2 inches apart.",
                "Bake for 9 to 11 minutes or until golden brown. Cool on baking sheets for 2 minutes, then transfer to wire racks.",
            ],
        },
        {
            "name": "Apple Pie",
            "ingredients": [
                "2 1/2 cups all-purpose flour",
                "1 tsp salt",
                "1 cup shortening, chilled",
                "6 to 8 tbsp ice water",
                "6 large apples, peeled and thinly sliced (about 8 cups)",
                "3/4 cup granulated sugar",
                "2 tbsp all-purpose flour",
                "3/4 tsp ground cinnamon",
                "1/4 tsp ground nutmeg",
                "1 tbsp butter, cut into small pieces",
            ],
            "directions": [
                "For the crust: In a large bowl, mix flour and salt. Cut in shortening until mixture resembles coarse crumbs. Sprinkle in ice water, one tablespoon at a time, tossing lightly with a fork until dough just holds together.",
                "Divide dough in half. On a lightly floured surface, roll each half into a circle about 12 inches across. Transfer one circle to a 9-inch pie plate.",
                "For the filling: Combine sliced apples, sugar, flour, cinnamon, and nutmeg in a large bowl. Toss to coat evenly.",
                "Pour apple mixture into the pie shell. Dot with butter pieces. Cover with the second pastry circle, trim, and crimp edges to seal. Cut several slits in the top crust to vent steam.",
                "Bake at 425°F for 15 minutes, then reduce heat to 350°F and bake for 40 to 45 more minutes until crust is golden and filling is bubbly.",
                "Cool on a wire rack for at least 1 hour before serving.",
            ],
        },
        {
            "name": "Banana Bread",
            "ingredients": [
                "1 1/2 cups all-purpose flour",
                "1 tsp baking soda",
                "1/4 tsp salt",
                "3 very ripe bananas, mashed",
                "3/4 cup granulated sugar",
                "1/3 cup butter, melted",
                "1 egg, beaten",
                "1 tsp vanilla extract",
                "1/2 cup chopped walnuts (optional)",
            ],
            "directions": [
                "Preheat oven to 350°F. Grease a 9x5 inch loaf pan.",
                "In a large bowl, stir together flour, baking soda, and salt.",
                "In a separate bowl, mix together the mashed bananas, sugar, melted butter, beaten egg, and vanilla extract until well combined.",
                "Pour the banana mixture into the flour mixture and stir just until moistened. Do not overmix. Fold in walnuts if using.",
                "Pour batter into the prepared loaf pan. Bake for 60 to 65 minutes, or until a toothpick inserted into the center comes out clean.",
                "Let the bread cool in the pan for 10 minutes before turning out onto a wire rack to cool completely.",
            ],
        },
        {
            "name": "Lemon Bars",
            "ingredients": [
                "For the crust:",
                "1 cup all-purpose flour",
                "1/4 cup powdered sugar",
                "1/2 cup (1 stick) butter, softened",
                "For the filling:",
                "2 large eggs",
                "1 cup granulated sugar",
                "2 tbsp all-purpose flour",
                "3 tbsp fresh lemon juice",
                "1 tsp grated lemon zest",
                "Powdered sugar for dusting",
            ],
            "directions": [
                "Preheat oven to 350°F. Grease an 8x8 inch baking pan.",
                "For the crust: Mix flour, powdered sugar, and softened butter until crumbly. Press evenly into the bottom of the prepared pan. Bake for 15 to 20 minutes until lightly golden.",
                "For the filling: While the crust bakes, whisk together the eggs, granulated sugar, flour, lemon juice, and lemon zest until smooth.",
                "Pour the lemon filling over the hot crust as soon as it comes out of the oven.",
                "Return to oven and bake for an additional 20 to 25 minutes, until the filling is set and does not jiggle.",
                "Allow to cool completely in the pan. Dust generously with powdered sugar before cutting into bars.",
            ],
        },
    ]

    for i, recipe in enumerate(recipes):
        # Recipe name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(recipe["name"])
        name_run.font.name = "Times New Roman"
        name_run.font.size = Pt(12)

        # Ingredients heading
        ing_para = doc.add_paragraph()
        ing_run = ing_para.add_run("Ingredients")
        ing_run.font.name = "Times New Roman"
        ing_run.font.size = Pt(12)

        # Each ingredient
        for ingredient in recipe["ingredients"]:
            ing_item = doc.add_paragraph()
            ing_item_run = ing_item.add_run(ingredient)
            ing_item_run.font.name = "Times New Roman"
            ing_item_run.font.size = Pt(12)

        # Directions heading
        dir_para = doc.add_paragraph()
        dir_run = dir_para.add_run("Directions")
        dir_run.font.name = "Times New Roman"
        dir_run.font.size = Pt(12)

        # Each direction
        for direction in recipe["directions"]:
            dir_item = doc.add_paragraph()
            dir_item_run = dir_item.add_run(direction)
            dir_item_run.font.name = "Times New Roman"
            dir_item_run.font.size = Pt(12)

        # Blank line between recipes (except after last)
        if i < len(recipes) - 1:
            blank = doc.add_paragraph()
            blank_run = blank.add_run("")
            blank_run.font.name = "Times New Roman"
            blank_run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
