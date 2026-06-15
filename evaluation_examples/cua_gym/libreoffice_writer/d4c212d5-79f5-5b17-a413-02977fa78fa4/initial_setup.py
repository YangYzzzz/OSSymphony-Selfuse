"""
Initial Setup: Apply 'Keep with next' to article headings in a legal contract
Task ID: writer_legal_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_040'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_no_keep(doc, text, level=1):
    """Add a heading and explicitly disable keep_with_next."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = False
    return h


def add_body_paragraphs(doc, paragraphs):
    """Add multiple body paragraphs with normal style."""
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)


def create_initial():
    doc = Document()

    # Page setup - standard legal letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "CONFIDENTIAL — Master Services Agreement"
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r0 = fp.add_run("Page ")
    r0.font.size = Pt(8)
    r1 = fp.add_run()
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    # ========================
    # Title Page
    # ========================
    title = doc.add_heading('MASTER SERVICES AGREEMENT', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(72)

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Between')
    run.font.size = Pt(14)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Meridian Global Technologies, Inc.')
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('("Provider")')
    run.font.size = Pt(12)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('and')
    run.font.size = Pt(14)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Pinnacle Financial Services, LLC')
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('("Client")')
    run.font.size = Pt(12)

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Contract Reference: MSA-2025-04781')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ========================
    # Table of Contents placeholder
    # ========================
    toc = doc.add_heading('TABLE OF CONTENTS', level=0)
    toc.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    toc_items = [
        "Article I — Definitions and Interpretation",
        "Article II — Scope of Services",
        "Article III — Term and Termination",
        "Article IV — Compensation and Payment Terms",
        "Article V — Intellectual Property Rights",
        "Article VI — Confidentiality Obligations",
        "Article VII — Representations and Warranties",
        "Article VIII — Indemnification",
        "Article IX — Limitation of Liability",
        "Article X — Data Protection and Privacy",
        "Article XI — Force Majeure",
        "Article XII — Insurance Requirements",
        "Article XIII — Compliance with Laws",
        "Article XIV — Dispute Resolution",
        "Article XV — Assignment and Subcontracting",
        "Article XVI — Notices",
        "Article XVII — Miscellaneous Provisions",
        "Article XVIII — Entire Agreement and Amendment",
        "Article XIX — Governing Law and Jurisdiction",
        "Article XX — Signatures and Execution",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ========================
    # PREAMBLE
    # ========================
    add_heading_no_keep(doc, 'PREAMBLE')
    add_body_paragraphs(doc, [
        'This Master Services Agreement ("Agreement") is entered into as of March 15, 2025 (the "Effective Date"), by and between Meridian Global Technologies, Inc., a Delaware corporation with its principal offices at 4500 Technology Parkway, Suite 800, San Jose, California 95134 ("Provider"), and Pinnacle Financial Services, LLC, a New York limited liability company with its principal offices at 200 Park Avenue, 35th Floor, New York, New York 10166 ("Client").',
        'WHEREAS, Provider is engaged in the business of providing enterprise technology solutions, software development services, cloud infrastructure management, cybersecurity consulting, and related professional services;',
        'WHEREAS, Client desires to engage Provider to perform certain technology services and deliver specific technology solutions as more particularly described in the Statements of Work attached hereto;',
        'WHEREAS, the parties wish to establish the general terms and conditions under which Provider shall provide such services to Client;',
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the parties agree as follows:',
    ])

    # ========================
    # Article I — Definitions
    # ========================
    add_heading_no_keep(doc, 'Article I — Definitions and Interpretation')

    add_body_paragraphs(doc, [
        '1.1 "Affiliate" means, with respect to any entity, any other entity that directly or indirectly controls, is controlled by, or is under common control with such entity, where "control" means the possession, directly or indirectly, of the power to direct or cause the direction of the management and policies of an entity, whether through the ownership of voting securities, by contract, or otherwise.',
        '1.2 "Business Day" means any day other than a Saturday, Sunday, or a day on which banks in New York, New York or San Jose, California are authorized or required by law to close.',
        '1.3 "Change Order" means a written document signed by both parties that modifies the scope, schedule, deliverables, or fees associated with a Statement of Work, substantially in the form attached hereto as Exhibit B.',
        '1.4 "Confidential Information" means all non-public information disclosed by one party to the other party, whether orally, in writing, or by inspection of tangible objects, including without limitation trade secrets, business plans, financial information, customer data, technical specifications, algorithms, source code, object code, designs, inventions, processes, techniques, marketing plans, and strategies.',
        '1.5 "Deliverables" means all work product, reports, documentation, software, tools, designs, data, and other materials created, developed, or delivered by Provider to Client in the performance of the Services under a Statement of Work.',
        '1.6 "Effective Date" means the date first written above in the preamble to this Agreement.',
        '1.7 "Fees" means the compensation payable to Provider for the Services as set forth in each Statement of Work, including any adjustments made through approved Change Orders.',
        '1.8 "Force Majeure Event" means any event beyond the reasonable control of the affected party, including but not limited to acts of God, natural disasters, epidemics, pandemics, government actions, war, terrorism, riots, embargoes, labor disputes not involving the affected party\'s employees, failures of third-party telecommunications or power supply, and cyberattacks perpetrated by third parties not affiliated with the affected party.',
        '1.9 "Intellectual Property" or "IP" means all patents, copyrights, trademarks, trade secrets, know-how, inventions, designs, algorithms, software code (both source and object), databases, documentation, and any other intellectual property rights, whether registered or unregistered.',
        '1.10 "Personal Data" means any information relating to an identified or identifiable natural person, as defined under applicable data protection laws, including the General Data Protection Regulation (EU) 2016/679, the California Consumer Privacy Act, and any successor legislation.',
        '1.11 "Project Manager" means the individual designated by each party to serve as the primary point of contact for day-to-day management and communication regarding the Services.',
        '1.12 "Services" means the professional services, consulting, development, implementation, maintenance, support, and any other services to be provided by Provider to Client as described in one or more Statements of Work executed under this Agreement.',
        '1.13 "Statement of Work" or "SOW" means a written document executed by both parties that describes the specific Services to be performed, the Deliverables to be provided, the timeline, the Fees, and any other terms applicable to a particular project or engagement, substantially in the form attached hereto as Exhibit A.',
    ])

    # ========================
    # Article II — Scope of Services
    # ========================
    add_heading_no_keep(doc, 'Article II — Scope of Services')

    add_body_paragraphs(doc, [
        '2.1 Provider shall perform the Services and deliver the Deliverables as described in each Statement of Work executed by the parties from time to time during the Term of this Agreement. Each SOW shall be deemed incorporated into and governed by the terms and conditions of this Agreement.',
        '2.2 In the event of any conflict or inconsistency between the terms of this Agreement and the terms of any SOW, the terms of this Agreement shall prevail unless the SOW expressly states that it is intended to supersede a specific provision of this Agreement, in which case the SOW shall control with respect to such specific provision only.',
        '2.3 Provider shall perform the Services in a professional and workmanlike manner, consistent with generally accepted industry standards and practices. Provider shall use commercially reasonable efforts to meet the milestones, deadlines, and performance standards specified in each SOW.',
        '2.4 Provider shall assign qualified personnel with appropriate skills, experience, and training to perform the Services. Client shall have the right to request the removal of any Provider personnel who, in Client\'s reasonable judgment, are not performing satisfactorily, and Provider shall promptly replace such personnel with individuals of comparable qualifications.',
        '2.5 Each party shall designate a Project Manager within five (5) Business Days of the execution of each SOW. The Project Managers shall meet regularly, no less than bi-weekly, to review progress, discuss issues, and coordinate activities under the applicable SOW.',
        '2.6 Any changes to the scope, schedule, or Fees under a SOW must be documented in a Change Order signed by both parties. Neither party shall be obligated to perform or pay for work outside the scope of an executed SOW or Change Order.',
        '2.7 Provider acknowledges that time is of the essence with respect to the performance of the Services and the delivery of the Deliverables in accordance with the schedule set forth in each SOW. If Provider anticipates any delay, Provider shall promptly notify Client in writing, specifying the nature of the delay, the expected duration, and any proposed mitigation measures.',
    ])

    # ========================
    # Article III — Term and Termination
    # ========================
    add_heading_no_keep(doc, 'Article III — Term and Termination')

    add_body_paragraphs(doc, [
        '3.1 This Agreement shall commence on the Effective Date and shall continue for an initial term of three (3) years (the "Initial Term"), unless earlier terminated in accordance with this Article III.',
        '3.2 Upon expiration of the Initial Term, this Agreement shall automatically renew for successive one (1) year periods (each, a "Renewal Term"), unless either party provides written notice of non-renewal to the other party at least ninety (90) days prior to the expiration of the then-current term.',
        '3.3 Either party may terminate this Agreement or any SOW for cause upon written notice if the other party: (a) materially breaches any provision of this Agreement or the applicable SOW and fails to cure such breach within thirty (30) days after receipt of written notice specifying the nature of the breach; or (b) becomes insolvent, files a petition for bankruptcy, has an involuntary petition filed against it that is not dismissed within sixty (60) days, or makes an assignment for the benefit of creditors.',
        '3.4 Client may terminate any SOW for convenience upon sixty (60) days\' prior written notice to Provider. In the event of such termination, Client shall pay Provider for all Services performed and Deliverables delivered through the effective date of termination, plus any reasonable, documented, non-cancellable costs incurred by Provider in connection with the terminated SOW.',
        '3.5 Upon termination or expiration of this Agreement for any reason: (a) each party shall promptly return or destroy all Confidential Information of the other party in its possession; (b) Provider shall deliver to Client all completed and in-progress Deliverables; (c) all licenses granted hereunder shall immediately terminate, except as expressly provided herein; and (d) the provisions of Articles V, VI, VII, VIII, IX, X, XIV, and XIX shall survive termination.',
        '3.6 Termination of this Agreement shall not release either party from any liability that has accrued as of the date of termination, nor shall it preclude either party from pursuing any rights or remedies available at law or in equity with respect to any breach of this Agreement occurring prior to termination.',
    ])

    # ========================
    # Article IV — Compensation and Payment
    # ========================
    add_heading_no_keep(doc, 'Article IV — Compensation and Payment Terms')

    add_body_paragraphs(doc, [
        '4.1 Client shall pay Provider the Fees as specified in each SOW. Unless otherwise stated in the applicable SOW, Fees shall be invoiced monthly in arrears based on the Services performed and Deliverables delivered during the preceding calendar month.',
        '4.2 All invoices shall be due and payable within thirty (30) days of receipt by Client. Invoices shall be submitted electronically to Client\'s accounts payable department at the email address specified in the applicable SOW, and shall include a detailed description of the Services performed, the hours worked (if applicable), and any supporting documentation reasonably required by Client.',
        '4.3 If Client disputes any portion of an invoice, Client shall notify Provider in writing within fifteen (15) days of receipt of the invoice, specifying the nature and basis of the dispute. The undisputed portion of the invoice shall remain due and payable in accordance with Section 4.2. The parties shall work in good faith to resolve any invoice dispute within thirty (30) days.',
        '4.4 Late payments shall accrue interest at the rate of one and one-half percent (1.5%) per month, or the maximum rate permitted by applicable law, whichever is less, calculated from the due date until the date of payment.',
        '4.5 All Fees are exclusive of applicable taxes, duties, and levies. Client shall be responsible for all sales, use, value-added, and similar taxes imposed on the Services, excluding taxes based on Provider\'s income. If Provider is required to collect or remit any such taxes, Provider shall add the applicable tax amount to its invoices.',
        '4.6 Provider shall maintain accurate and complete time records and expense reports in connection with the Services. Client shall have the right, upon thirty (30) days\' prior written notice, to audit Provider\'s books and records related to the Fees and expenses charged under this Agreement, not more than once per calendar year.',
        '4.7 In the event that any governmental authority imposes new taxes, duties, or assessments applicable to the Services after the Effective Date, the parties shall negotiate in good faith to equitably allocate such additional costs. If the parties are unable to reach agreement within thirty (30) days, either party may terminate the affected SOW upon thirty (30) days\' written notice.',
    ])

    # ========================
    # Article V — Intellectual Property
    # ========================
    add_heading_no_keep(doc, 'Article V — Intellectual Property Rights')

    add_body_paragraphs(doc, [
        '5.1 Pre-Existing IP. Each party shall retain all right, title, and interest in and to its pre-existing Intellectual Property. Neither party grants the other any rights in its pre-existing IP except as expressly set forth in this Agreement.',
        '5.2 Work Product Ownership. Subject to Section 5.1 and Section 5.3, all Deliverables and work product created by Provider specifically for Client in the performance of the Services under a SOW shall be considered "works made for hire" to the extent permitted by applicable law. To the extent any Deliverable does not qualify as a work made for hire, Provider hereby irrevocably assigns to Client all right, title, and interest in and to such Deliverable, including all Intellectual Property rights therein.',
        '5.3 Provider Tools and Methodologies. Notwithstanding Section 5.2, Provider shall retain all right, title, and interest in and to: (a) any tools, frameworks, libraries, methodologies, processes, techniques, and know-how that were developed by Provider independently of the Services or that are of general applicability ("Provider Tools"); and (b) any improvements or modifications to Provider Tools made during the performance of the Services. Provider hereby grants Client a non-exclusive, perpetual, irrevocable, royalty-free, worldwide license to use any Provider Tools incorporated into the Deliverables, solely in connection with Client\'s use of the Deliverables.',
        '5.4 Third-Party Components. If any Deliverable incorporates third-party software or components, Provider shall: (a) identify such third-party components in the applicable SOW; (b) ensure that the license terms applicable to such components permit their use as contemplated under this Agreement; and (c) provide Client with copies of all applicable third-party license agreements.',
        '5.5 Provider shall obtain from each of its employees and subcontractors engaged in the performance of the Services a written assignment of all Intellectual Property rights in the work product sufficient to effectuate the assignments and licenses granted to Client under this Article V.',
    ])

    # ========================
    # Article VI — Confidentiality
    # ========================
    add_heading_no_keep(doc, 'Article VI — Confidentiality Obligations')

    add_body_paragraphs(doc, [
        '6.1 Each party (the "Receiving Party") agrees to hold in strict confidence all Confidential Information received from the other party (the "Disclosing Party") and to use such Confidential Information solely for the purpose of performing its obligations or exercising its rights under this Agreement.',
        '6.2 The Receiving Party shall not disclose Confidential Information to any third party without the prior written consent of the Disclosing Party, except that the Receiving Party may disclose Confidential Information to its employees, contractors, advisors, and Affiliates who have a need to know such information for purposes consistent with this Agreement, provided that such persons are bound by confidentiality obligations no less restrictive than those set forth herein.',
        '6.3 The confidentiality obligations set forth in this Article VI shall not apply to information that: (a) is or becomes publicly available other than through a breach of this Agreement; (b) was known to the Receiving Party prior to disclosure by the Disclosing Party, as evidenced by written records; (c) is independently developed by the Receiving Party without use of or reference to the Confidential Information; or (d) is rightfully received by the Receiving Party from a third party without restriction on disclosure.',
        '6.4 If the Receiving Party is compelled by law, regulation, or legal process to disclose Confidential Information, it shall provide the Disclosing Party with prompt written notice thereof (to the extent permitted by law) and shall cooperate with the Disclosing Party in seeking a protective order or other appropriate remedy to limit such disclosure.',
        '6.5 Upon termination or expiration of this Agreement, or upon the Disclosing Party\'s written request, the Receiving Party shall promptly return or destroy all Confidential Information in its possession, including all copies, notes, summaries, and extracts thereof, and shall certify in writing that it has done so.',
        '6.6 The confidentiality obligations under this Article VI shall survive the termination or expiration of this Agreement for a period of five (5) years; provided, however, that the obligations with respect to trade secrets shall continue for as long as such information qualifies as a trade secret under applicable law.',
    ])

    # ========================
    # Article VII — Representations and Warranties
    # ========================
    add_heading_no_keep(doc, 'Article VII — Representations and Warranties')

    add_body_paragraphs(doc, [
        '7.1 Mutual Representations. Each party represents and warrants to the other that: (a) it is duly organized, validly existing, and in good standing under the laws of its jurisdiction of formation; (b) it has full power and authority to enter into this Agreement and to perform its obligations hereunder; (c) the execution, delivery, and performance of this Agreement have been duly authorized by all necessary corporate action; and (d) this Agreement constitutes a legal, valid, and binding obligation enforceable in accordance with its terms.',
        '7.2 Provider Warranties. Provider represents and warrants that: (a) the Services shall be performed in a professional and workmanlike manner by qualified personnel; (b) the Deliverables shall conform in all material respects to the specifications set forth in the applicable SOW; (c) the Deliverables shall be free from material defects for a period of ninety (90) days following acceptance by Client (the "Warranty Period"); (d) the Services and Deliverables shall not infringe, misappropriate, or otherwise violate any Intellectual Property rights of any third party; and (e) Provider shall comply with all applicable laws, regulations, and industry standards in the performance of the Services.',
        '7.3 If any Deliverable fails to conform to the warranties set forth in Section 7.2 during the Warranty Period, Provider shall, at its own expense, promptly correct such non-conformity. If Provider is unable to correct the non-conformity within thirty (30) days after receipt of written notice from Client, Client may, at its option: (a) require Provider to re-perform the applicable Services at no additional charge; or (b) terminate the applicable SOW and receive a refund of any Fees paid for the non-conforming Deliverable.',
        '7.4 EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, NEITHER PARTY MAKES ANY WARRANTIES, WHETHER EXPRESS, IMPLIED, STATUTORY, OR OTHERWISE, INCLUDING ANY IMPLIED WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, TITLE, OR NON-INFRINGEMENT.',
    ])

    # ========================
    # Article VIII — Indemnification
    # ========================
    add_heading_no_keep(doc, 'Article VIII — Indemnification')

    add_body_paragraphs(doc, [
        '8.1 Provider Indemnification. Provider shall defend, indemnify, and hold harmless Client and its Affiliates, and their respective officers, directors, employees, agents, successors, and assigns (collectively, the "Client Indemnitees") from and against any and all losses, damages, liabilities, claims, actions, judgments, settlements, interest, penalties, fines, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to: (a) any breach by Provider of its representations, warranties, or obligations under this Agreement; (b) any claim that the Services or Deliverables infringe, misappropriate, or otherwise violate any Intellectual Property rights of any third party; (c) Provider\'s negligence or willful misconduct in the performance of the Services; or (d) any violation of applicable law by Provider in connection with this Agreement.',
        '8.2 Client Indemnification. Client shall defend, indemnify, and hold harmless Provider and its Affiliates, and their respective officers, directors, employees, agents, successors, and assigns (collectively, the "Provider Indemnitees") from and against any and all losses, damages, liabilities, claims, actions, judgments, settlements, interest, penalties, fines, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to: (a) any breach by Client of its representations, warranties, or obligations under this Agreement; (b) Client\'s negligence or willful misconduct; or (c) any claim arising from Client\'s use of the Deliverables in a manner not contemplated by this Agreement.',
        '8.3 Indemnification Procedure. The party seeking indemnification (the "Indemnified Party") shall: (a) promptly notify the indemnifying party (the "Indemnifying Party") in writing of any claim for which indemnification is sought; (b) grant the Indemnifying Party sole control of the defense and settlement of such claim; and (c) provide reasonable cooperation and assistance at the Indemnifying Party\'s expense. The Indemnifying Party shall not settle any claim without the prior written consent of the Indemnified Party if such settlement would impose any obligation on the Indemnified Party or adversely affect its rights.',
        '8.4 If any Deliverable becomes, or in Provider\'s reasonable opinion is likely to become, the subject of an infringement claim, Provider may, at its sole expense and option: (a) procure for Client the right to continue using the Deliverable; (b) modify the Deliverable to make it non-infringing while maintaining substantially equivalent functionality; (c) replace the Deliverable with a non-infringing alternative of substantially equivalent functionality; or (d) if none of the foregoing options are commercially practicable, terminate the applicable SOW and refund to Client the Fees paid for the infringing Deliverable.',
    ])

    # ========================
    # Article IX — Limitation of Liability
    # ========================
    add_heading_no_keep(doc, 'Article IX — Limitation of Liability')

    add_body_paragraphs(doc, [
        '9.1 EXCEPT FOR (A) EACH PARTY\'S INDEMNIFICATION OBLIGATIONS UNDER ARTICLE VIII, (B) EITHER PARTY\'S BREACH OF ITS CONFIDENTIALITY OBLIGATIONS UNDER ARTICLE VI, (C) PROVIDER\'S BREACH OF ITS INTELLECTUAL PROPERTY OBLIGATIONS UNDER ARTICLE V, AND (D) LIABILITY ARISING FROM A PARTY\'S GROSS NEGLIGENCE OR WILLFUL MISCONDUCT, NEITHER PARTY SHALL BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, PUNITIVE, OR EXEMPLARY DAMAGES, INCLUDING BUT NOT LIMITED TO DAMAGES FOR LOSS OF PROFITS, GOODWILL, USE, DATA, OR OTHER INTANGIBLE LOSSES, REGARDLESS OF WHETHER SUCH PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.',
        '9.2 SUBJECT TO THE EXCEPTIONS SET FORTH IN SECTION 9.1, EACH PARTY\'S TOTAL AGGREGATE LIABILITY ARISING OUT OF OR RELATED TO THIS AGREEMENT, WHETHER IN CONTRACT, TORT, OR OTHERWISE, SHALL NOT EXCEED THE GREATER OF: (A) THE TOTAL FEES PAID OR PAYABLE BY CLIENT TO PROVIDER UNDER THIS AGREEMENT DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM; OR (B) TWO MILLION DOLLARS ($2,000,000).',
        '9.3 The limitations set forth in this Article IX shall apply to the maximum extent permitted by applicable law, and shall not be affected by the failure of any limited remedy to achieve its essential purpose. Each party acknowledges that the other party has entered into this Agreement in reliance upon the limitations of liability set forth herein, and that such limitations represent a fair and reasonable allocation of risk between the parties.',
    ])

    # ========================
    # Article X — Data Protection
    # ========================
    add_heading_no_keep(doc, 'Article X — Data Protection and Privacy')

    add_body_paragraphs(doc, [
        '10.1 To the extent that Provider processes Personal Data on behalf of Client in the performance of the Services, Provider shall process such Personal Data only in accordance with Client\'s documented instructions and applicable data protection laws.',
        '10.2 Provider shall implement and maintain appropriate technical and organizational measures to protect Personal Data against unauthorized access, alteration, disclosure, or destruction, including but not limited to: (a) encryption of Personal Data in transit and at rest; (b) regular security assessments and penetration testing; (c) access controls limiting access to Personal Data to authorized personnel; (d) incident response and breach notification procedures; and (e) regular backup and disaster recovery procedures.',
        '10.3 Provider shall not transfer Personal Data to any country or territory outside the European Economic Area or the United States without the prior written consent of Client and without ensuring that adequate data protection safeguards are in place, as required by applicable data protection laws.',
        '10.4 In the event of a security breach involving Personal Data, Provider shall: (a) notify Client without undue delay, and in any event within forty-eight (48) hours of becoming aware of the breach; (b) provide Client with sufficient information to enable Client to comply with its own notification obligations under applicable law; (c) cooperate with Client in investigating and remediating the breach; and (d) take all reasonable steps to mitigate the effects of the breach and prevent recurrence.',
        '10.5 Upon termination or expiration of this Agreement, Provider shall, at Client\'s election, return or securely delete all Personal Data in its possession within thirty (30) days, and shall certify in writing that it has done so. Provider may retain copies of Personal Data only to the extent required by applicable law, subject to continued compliance with this Article X.',
        '10.6 Client shall have the right to audit Provider\'s compliance with this Article X, upon thirty (30) days\' prior written notice, not more than once per year. Provider shall provide reasonable cooperation and access to its facilities, systems, and records as necessary to facilitate such audit.',
    ])

    # ========================
    # Article XI — Force Majeure
    # ========================
    add_heading_no_keep(doc, 'Article XI — Force Majeure')

    add_body_paragraphs(doc, [
        '11.1 Neither party shall be liable for any failure or delay in the performance of its obligations under this Agreement (other than payment obligations) to the extent such failure or delay results from a Force Majeure Event.',
        '11.2 The affected party shall provide prompt written notice to the other party of the Force Majeure Event, including a description of the event, the anticipated duration, and the obligations affected. The affected party shall use commercially reasonable efforts to mitigate the impact of the Force Majeure Event and resume performance as soon as practicable.',
        '11.3 If a Force Majeure Event continues for more than sixty (60) consecutive days, either party may terminate the affected SOW, or this Agreement if the Force Majeure Event affects all active SOWs, upon written notice to the other party. In such event, the provisions of Section 3.5 shall apply.',
        '11.4 During the period of a Force Majeure Event, the obligations of the affected party under this Agreement shall be suspended to the extent necessitated by the event, and the time for performance of such obligations shall be extended by a period equal to the duration of the Force Majeure Event.',
    ])

    # ========================
    # Article XII — Insurance
    # ========================
    add_heading_no_keep(doc, 'Article XII — Insurance Requirements')

    add_body_paragraphs(doc, [
        '12.1 Provider shall maintain, at its own expense, the following insurance coverage during the Term of this Agreement and for a period of two (2) years following termination or expiration: (a) commercial general liability insurance with limits of not less than $2,000,000 per occurrence and $5,000,000 in the aggregate; (b) professional liability (errors and omissions) insurance with limits of not less than $5,000,000 per claim and $10,000,000 in the aggregate; (c) workers\' compensation insurance as required by applicable law; (d) employer\'s liability insurance with limits of not less than $1,000,000 per occurrence; and (e) cyber liability insurance with limits of not less than $5,000,000 per claim.',
        '12.2 All insurance policies required under this Article XII shall: (a) be issued by insurance companies with an A.M. Best rating of A- VII or better; (b) name Client as an additional insured on the commercial general liability and cyber liability policies; (c) include a waiver of subrogation in favor of Client; and (d) provide that such policies shall not be cancelled, materially modified, or allowed to expire without at least thirty (30) days\' prior written notice to Client.',
        '12.3 Upon Client\'s request, Provider shall furnish Client with certificates of insurance evidencing the coverage required under this Article XII, including endorsements naming Client as an additional insured. Provider shall provide updated certificates annually and upon any material change in coverage.',
    ])

    # ========================
    # Article XIII — Compliance
    # ========================
    add_heading_no_keep(doc, 'Article XIII — Compliance with Laws')

    add_body_paragraphs(doc, [
        '13.1 Each party shall comply with all applicable federal, state, local, and international laws, regulations, rules, and orders in the performance of its obligations under this Agreement, including but not limited to: (a) anti-bribery and anti-corruption laws, including the U.S. Foreign Corrupt Practices Act and the U.K. Bribery Act 2010; (b) export control and economic sanctions laws; (c) employment and labor laws; (d) environmental laws; and (e) data protection and privacy laws.',
        '13.2 Provider represents and warrants that neither it nor any of its officers, directors, employees, agents, or subcontractors has, directly or indirectly, offered, promised, given, or authorized the giving of any money, gift, or anything of value to any government official, political party, or candidate for political office for the purpose of influencing any official act or decision, or obtaining or retaining business.',
        '13.3 Provider shall maintain adequate internal controls and compliance programs to ensure compliance with applicable laws and regulations. Provider shall promptly notify Client in writing of any material violation or potential violation of applicable law that may affect the Services or Client\'s interests.',
        '13.4 If any change in applicable law or regulation materially affects the Services or the ability of either party to perform its obligations under this Agreement, the parties shall negotiate in good faith to amend this Agreement or the affected SOW to address such change.',
    ])

    # ========================
    # Article XIV — Dispute Resolution
    # ========================
    add_heading_no_keep(doc, 'Article XIV — Dispute Resolution')

    add_body_paragraphs(doc, [
        '14.1 The parties shall attempt in good faith to resolve any dispute, controversy, or claim arising out of or relating to this Agreement or the breach, termination, or validity thereof (a "Dispute") through informal negotiation. Either party may initiate informal negotiations by delivering written notice to the other party describing the nature of the Dispute and the relief sought.',
        '14.2 If the Dispute is not resolved within thirty (30) days of the initial notice, either party may escalate the Dispute to the senior executives of each party designated for such purpose, who shall meet (in person or via videoconference) within fifteen (15) days of escalation and attempt in good faith to resolve the Dispute.',
        '14.3 If the Dispute is not resolved within sixty (60) days of the initial notice through the procedures set forth in Sections 14.1 and 14.2, either party may submit the Dispute to binding arbitration administered by the American Arbitration Association ("AAA") in accordance with its Commercial Arbitration Rules. The arbitration shall be conducted by a panel of three (3) arbitrators, with each party selecting one arbitrator and the two party-selected arbitrators selecting the third. The arbitration shall be held in New York, New York.',
        '14.4 The arbitrators shall have the authority to award any remedy or relief that a court of competent jurisdiction could grant, including specific performance, injunctive relief, and compensatory damages, but shall not have the authority to award punitive, exemplary, or consequential damages beyond those permitted under Article IX. The arbitrators\' decision shall be final and binding, and judgment thereon may be entered in any court of competent jurisdiction.',
        '14.5 Notwithstanding the foregoing, either party may seek provisional or injunctive relief from a court of competent jurisdiction where necessary to protect its Confidential Information, Intellectual Property rights, or other proprietary interests, pending resolution of the Dispute through arbitration.',
        '14.6 Each party shall bear its own costs and expenses in connection with any dispute resolution proceedings under this Article XIV, including attorneys\' fees, unless the arbitrators determine that one party\'s position was frivolous or maintained in bad faith, in which case the arbitrators may award costs and reasonable attorneys\' fees to the prevailing party.',
    ])

    # ========================
    # Article XV — Assignment
    # ========================
    add_heading_no_keep(doc, 'Article XV — Assignment and Subcontracting')

    add_body_paragraphs(doc, [
        '15.1 Neither party may assign this Agreement or any rights or obligations hereunder without the prior written consent of the other party, which consent shall not be unreasonably withheld, conditioned, or delayed; provided, however, that either party may assign this Agreement without consent to an Affiliate or in connection with a merger, acquisition, corporate reorganization, or sale of all or substantially all of its assets, provided that the assignee agrees in writing to be bound by the terms of this Agreement.',
        '15.2 Provider may subcontract any portion of the Services to qualified third-party subcontractors, subject to the following conditions: (a) Provider shall notify Client in writing prior to engaging any subcontractor; (b) Provider shall ensure that each subcontractor is bound by confidentiality, intellectual property, and data protection obligations no less restrictive than those set forth in this Agreement; (c) Provider shall remain fully responsible and liable for the acts, omissions, and performance of its subcontractors; and (d) Client shall have the right to object to any proposed subcontractor on reasonable grounds.',
        '15.3 Any purported assignment or subcontracting in violation of this Article XV shall be null and void.',
    ])

    # ========================
    # Article XVI — Notices
    # ========================
    add_heading_no_keep(doc, 'Article XVI — Notices')

    add_body_paragraphs(doc, [
        '16.1 All notices, requests, demands, and other communications required or permitted under this Agreement shall be in writing and shall be deemed duly given: (a) upon personal delivery; (b) one (1) Business Day after deposit with a nationally recognized overnight courier service; (c) three (3) Business Days after deposit in the United States mail, postage prepaid, certified or registered mail, return receipt requested; or (d) upon confirmation of receipt when sent by email to the addresses specified below, provided that a copy is also sent by one of the methods specified in clauses (a) through (c) within two (2) Business Days.',
        '16.2 Notices shall be sent to the following addresses, or to such other address as a party may designate by written notice to the other party:',
        'If to Provider: Meridian Global Technologies, Inc., Attn: General Counsel, 4500 Technology Parkway, Suite 800, San Jose, California 95134. Email: legal@meridianglobaltech.com.',
        'If to Client: Pinnacle Financial Services, LLC, Attn: Chief Legal Officer, 200 Park Avenue, 35th Floor, New York, New York 10166. Email: legal@pinnaclefinancial.com.',
    ])

    # ========================
    # Article XVII — Miscellaneous
    # ========================
    add_heading_no_keep(doc, 'Article XVII — Miscellaneous Provisions')

    add_body_paragraphs(doc, [
        '17.1 Independent Contractor. The relationship between the parties is that of independent contractors. Nothing in this Agreement shall be construed to create a partnership, joint venture, agency, or employment relationship between the parties. Neither party shall have the authority to bind the other or to incur any obligation on behalf of the other without the other party\'s prior written consent.',
        '17.2 Severability. If any provision of this Agreement is held to be invalid, illegal, or unenforceable by a court of competent jurisdiction, such provision shall be modified to the minimum extent necessary to make it valid, legal, and enforceable, and the remaining provisions shall continue in full force and effect.',
        '17.3 Waiver. The failure of either party to enforce any provision of this Agreement shall not constitute a waiver of that party\'s right to enforce that provision or any other provision at a later time. Any waiver must be in writing and signed by the waiving party.',
        '17.4 Cumulative Remedies. The rights and remedies provided in this Agreement are cumulative and are not exclusive of any other rights or remedies available at law or in equity.',
        '17.5 Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original, and all of which together shall constitute one and the same instrument. Electronic signatures and PDF copies of signatures shall be deemed original signatures for all purposes.',
        '17.6 Publicity. Neither party shall use the other party\'s name, trademarks, logos, or other proprietary identifiers in any press release, marketing material, or public communication without the prior written consent of the other party, except as required by applicable law or regulation.',
    ])

    # ========================
    # Article XVIII — Entire Agreement
    # ========================
    add_heading_no_keep(doc, 'Article XVIII — Entire Agreement and Amendment')

    add_body_paragraphs(doc, [
        '18.1 This Agreement, together with all SOWs, Change Orders, and Exhibits attached hereto, constitutes the entire agreement between the parties with respect to the subject matter hereof and supersedes all prior and contemporaneous agreements, understandings, negotiations, and discussions, whether oral or written, between the parties.',
        '18.2 No amendment, modification, or supplement to this Agreement shall be effective unless it is in writing and signed by authorized representatives of both parties. For the avoidance of doubt, no provision of this Agreement may be amended by course of dealing, usage of trade, or course of performance.',
        '18.3 In the event of any conflict between this Agreement and any Exhibit, SOW, or Change Order, the order of precedence shall be: (a) the applicable Change Order; (b) the applicable SOW; (c) the Exhibits; and (d) this Agreement.',
    ])

    # ========================
    # Article XIX — Governing Law
    # ========================
    add_heading_no_keep(doc, 'Article XIX — Governing Law and Jurisdiction')

    add_body_paragraphs(doc, [
        '19.1 This Agreement shall be governed by and construed in accordance with the laws of the State of New York, without regard to its conflict of laws principles.',
        '19.2 Subject to the arbitration provisions set forth in Article XIV, each party irrevocably submits to the exclusive jurisdiction of the federal and state courts located in the Borough of Manhattan, New York, New York, for any action, suit, or proceeding arising out of or relating to this Agreement that is not subject to arbitration.',
        '19.3 Each party hereby waives, to the fullest extent permitted by applicable law, any right to trial by jury in any action, proceeding, or counterclaim arising out of or relating to this Agreement.',
        '19.4 The United Nations Convention on Contracts for the International Sale of Goods shall not apply to this Agreement.',
    ])

    # ========================
    # Article XX — Signatures
    # ========================
    add_heading_no_keep(doc, 'Article XX — Signatures and Execution')

    add_body_paragraphs(doc, [
        'IN WITNESS WHEREOF, the parties have caused this Agreement to be executed by their duly authorized representatives as of the Effective Date.',
    ])

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Signature blocks
    p = doc.add_paragraph('MERIDIAN GLOBAL TECHNOLOGIES, INC.')
    p.runs[0].bold = True
    doc.add_paragraph('')
    doc.add_paragraph('_____________________________________')
    p = doc.add_paragraph('Name: Dr. Elena Vasquez')
    doc.add_paragraph('Title: Chief Executive Officer')
    doc.add_paragraph('Date: March 15, 2025')

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph('PINNACLE FINANCIAL SERVICES, LLC')
    p.runs[0].bold = True
    doc.add_paragraph('')
    doc.add_paragraph('_____________________________________')
    p = doc.add_paragraph('Name: Robert W. Hamilton III')
    doc.add_paragraph('Title: Managing Director')
    doc.add_paragraph('Date: March 15, 2025')

    doc.add_page_break()

    # ========================
    # Exhibit A
    # ========================
    add_heading_no_keep(doc, 'Exhibit A — Form of Statement of Work')

    add_body_paragraphs(doc, [
        'STATEMENT OF WORK No. [___]',
        'Under the Master Services Agreement dated March 15, 2025, between Meridian Global Technologies, Inc. ("Provider") and Pinnacle Financial Services, LLC ("Client").',
        '',
        '1. PROJECT DESCRIPTION: [Description of the specific project and objectives]',
        '2. SCOPE OF SERVICES: [Detailed description of Services to be performed]',
        '3. DELIVERABLES: [List of specific Deliverables with acceptance criteria]',
        '4. TIMELINE AND MILESTONES: [Project schedule with key milestones and deadlines]',
        '5. FEES AND PAYMENT: [Fee structure, rate cards, expense policy, payment schedule]',
        '6. PROJECT MANAGERS: Provider: [Name/Title] | Client: [Name/Title]',
        '7. ASSUMPTIONS AND DEPENDENCIES: [Key assumptions and external dependencies]',
        '8. ACCEPTANCE CRITERIA: [Specific criteria for acceptance of each Deliverable]',
        '',
        'AGREED AND ACCEPTED:',
        '',
        'Provider: _________________ Date: _________',
        'Client: _________________ Date: _________',
    ])

    doc.add_page_break()

    # ========================
    # Exhibit B
    # ========================
    add_heading_no_keep(doc, 'Exhibit B — Form of Change Order')

    add_body_paragraphs(doc, [
        'CHANGE ORDER No. [___]',
        'Under Statement of Work No. [___] dated [___], pursuant to the Master Services Agreement dated March 15, 2025.',
        '',
        '1. DESCRIPTION OF CHANGE: [Detailed description of the proposed change]',
        '2. REASON FOR CHANGE: [Business justification for the change]',
        '3. IMPACT ON SCOPE: [Changes to the scope of Services and Deliverables]',
        '4. IMPACT ON SCHEDULE: [Changes to milestones and delivery dates]',
        '5. IMPACT ON FEES: [Additional or reduced Fees resulting from the change]',
        '6. IMPACT ON RESOURCES: [Changes to resource allocation or staffing]',
        '',
        'This Change Order shall become effective upon execution by both parties and shall be incorporated into and governed by the terms of the Agreement and the applicable SOW.',
        '',
        'AGREED AND ACCEPTED:',
        '',
        'Provider: _________________ Date: _________',
        'Client: _________________ Date: _________',
    ])

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
