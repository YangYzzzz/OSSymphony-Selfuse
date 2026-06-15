"""
Initial Setup: Create a report document with a 3-level TOC, all entries in black.
Task ID: writer_mt_089
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_089'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_field(doc):
    """Insert a TOC field code that covers 3 heading levels."""
    para = doc.add_paragraph()
    run = para.add_run()
    r_element = run._element

    # Begin field
    fld_begin = r_element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r_element.append(fld_begin)

    run2 = para.add_run()
    instr = run2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = para.add_run()
    fld_sep = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run3._element.append(fld_sep)

    # Add static TOC entries so the TOC is visible without updating
    return para


def add_toc_entry(doc, text, level, style_name):
    """Add a static TOC entry paragraph with proper style."""
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # All black initially
    return para


def end_toc_field(doc, toc_start_para):
    """End the TOC field - add to last TOC entry paragraph."""
    # We need to close the field in a paragraph after all TOC entries
    para = doc.add_paragraph()
    run = para.add_run()
    fld_end = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run._element.append(fld_end)
    return para


def ensure_toc_styles(doc):
    """Ensure TOC Heading / Contents styles exist with black color."""
    styles = doc.styles
    for level in range(1, 4):
        style_name = f'toc {level}'
        try:
            s = styles[style_name]
        except KeyError:
            # Create TOC style if missing
            s = styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
            s.font.size = Pt(12 - level)  # 11, 10, 9
        # Make sure font color is black
        s.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # Set indentation based on level
        s.paragraph_format.left_indent = Inches(0.2 * (level - 1))
        s.paragraph_format.space_after = Pt(2)
        s.paragraph_format.space_before = Pt(2)


def create_initial():
    doc = Document()

    # Set up default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Title
    title = doc.add_heading('Quarterly Business Performance Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Fiscal Year 2025 - Q1 through Q4 Analysis')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Ensure TOC styles exist
    ensure_toc_styles(doc)

    # --- Table of Contents (static entries) ---
    toc_heading = doc.add_heading('Table of Contents', level=1)

    # TOC field start
    toc_start = add_toc_field(doc)

    # Heading 1 entries (5 items) - TOC level 1
    h1_entries = [
        'Executive Summary',
        'Financial Performance Overview',
        'Operational Metrics and KPIs',
        'Market Analysis and Competitive Landscape',
        'Strategic Recommendations and Outlook',
    ]

    # Heading 2 entries (12 items) - TOC level 2
    h2_entries = [
        'Revenue Breakdown by Region',
        'Cost Structure Analysis',
        'Profit Margin Trends',
        'Customer Acquisition Metrics',
        'Employee Productivity Index',
        'Supply Chain Efficiency',
        'Digital Transformation Progress',
        'Risk Assessment Summary',
        'Regulatory Compliance Status',
        'Technology Infrastructure Updates',
        'Talent Management Overview',
        'Sustainability Initiatives',
    ]

    # Heading 3 entries (8 items) - TOC level 3
    h3_entries = [
        'North America Revenue Detail',
        'European Market Expansion',
        'Asia-Pacific Growth Indicators',
        'Raw Material Cost Trends',
        'New Customer Onboarding Rate',
        'Server Uptime and Performance',
        'Carbon Footprint Reduction',
        'Diversity and Inclusion Metrics',
    ]

    # Build TOC entries in a realistic interleaved order
    toc_structure = [
        (h1_entries[0], 1),  # Executive Summary
        (h1_entries[1], 1),  # Financial Performance Overview
        (h2_entries[0], 2),  # Revenue Breakdown by Region
        (h3_entries[0], 3),  # North America Revenue Detail
        (h3_entries[1], 3),  # European Market Expansion
        (h3_entries[2], 3),  # Asia-Pacific Growth Indicators
        (h2_entries[1], 2),  # Cost Structure Analysis
        (h3_entries[3], 3),  # Raw Material Cost Trends
        (h2_entries[2], 2),  # Profit Margin Trends
        (h1_entries[2], 1),  # Operational Metrics and KPIs
        (h2_entries[3], 2),  # Customer Acquisition Metrics
        (h3_entries[4], 3),  # New Customer Onboarding Rate
        (h2_entries[4], 2),  # Employee Productivity Index
        (h2_entries[5], 2),  # Supply Chain Efficiency
        (h1_entries[3], 1),  # Market Analysis and Competitive Landscape
        (h2_entries[6], 2),  # Digital Transformation Progress
        (h2_entries[7], 2),  # Risk Assessment Summary
        (h2_entries[8], 2),  # Regulatory Compliance Status
        (h3_entries[5], 3),  # Server Uptime and Performance
        (h1_entries[4], 1),  # Strategic Recommendations and Outlook
        (h2_entries[9], 2),  # Technology Infrastructure Updates
        (h2_entries[10], 2),  # Talent Management Overview
        (h2_entries[11], 2),  # Sustainability Initiatives
        (h3_entries[6], 3),  # Carbon Footprint Reduction
        (h3_entries[7], 3),  # Diversity and Inclusion Metrics
    ]

    for entry_text, level in toc_structure:
        style_name = f'toc {level}'
        add_toc_entry(doc, entry_text, level, style_name)

    # End TOC field
    end_toc_field(doc, toc_start)

    doc.add_paragraph()  # spacer

    # --- Document body with actual headings and content ---

    # Section 1: Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive analysis of the company\'s performance '
        'across all four quarters of fiscal year 2025. Overall revenue grew by 14.2% '
        'year-over-year, reaching $847.3 million. Operating margins improved by 2.1 '
        'percentage points, driven by operational efficiencies and strategic cost management.'
    )
    doc.add_paragraph(
        'Key highlights include the successful expansion into three new markets, '
        'a 23% improvement in customer retention rates, and the completion of our '
        'digital transformation initiative Phase II. Employee satisfaction scores '
        'reached an all-time high of 4.6 out of 5.0.'
    )

    # Section 2: Financial Performance Overview
    doc.add_heading('Financial Performance Overview', level=1)
    doc.add_paragraph(
        'The financial results for FY2025 demonstrate sustained growth across all '
        'major business segments. Total revenue of $847.3M exceeded our target of '
        '$810M by 4.6%. Gross profit margin stood at 62.8%, up from 60.7% in the prior year.'
    )

    doc.add_heading('Revenue Breakdown by Region', level=2)
    doc.add_paragraph(
        'North America contributed $412.5M (48.7%), Europe $238.1M (28.1%), '
        'and Asia-Pacific $196.7M (23.2%). The fastest growth was observed in '
        'the Asia-Pacific segment at 22.4% year-over-year.'
    )

    doc.add_heading('North America Revenue Detail', level=3)
    doc.add_paragraph(
        'The US market accounted for $358.2M while Canada contributed $54.3M. '
        'Enterprise segment sales grew by 18.3%, offsetting a modest decline in '
        'consumer sales. Key enterprise wins included contracts with Meridian Healthcare, '
        'Atlas Financial Group, and Pinnacle Manufacturing.'
    )

    doc.add_heading('European Market Expansion', level=3)
    doc.add_paragraph(
        'The EU region saw revenue increase by 16.7%, with Germany, France, and the UK '
        'as the top three contributors. The new Berlin office became fully operational in Q2, '
        'adding 45 sales and support staff. The Nordic expansion pilot exceeded targets by 12%.'
    )

    doc.add_heading('Asia-Pacific Growth Indicators', level=3)
    doc.add_paragraph(
        'Japan remained the largest APAC market at $78.4M, followed by Australia at $52.1M '
        'and Singapore at $31.8M. India operations launched in Q3, generating $12.2M in the '
        'first two quarters with a pipeline valued at $45M for FY2026.'
    )

    doc.add_heading('Cost Structure Analysis', level=2)
    doc.add_paragraph(
        'Total operating expenses were $315.2M, representing 37.2% of revenue. This marks '
        'a 1.3 percentage point improvement over the prior year. Major cost categories include '
        'personnel ($178.4M), technology ($62.1M), facilities ($38.7M), and marketing ($36.0M).'
    )

    doc.add_heading('Raw Material Cost Trends', level=3)
    doc.add_paragraph(
        'Component costs decreased by 8.4% in H2 due to renegotiated supplier contracts '
        'and increased procurement volume discounts. The new strategic sourcing initiative '
        'is projected to save an additional $12M annually starting in FY2026.'
    )

    doc.add_heading('Profit Margin Trends', level=2)
    doc.add_paragraph(
        'Quarterly profit margins showed consistent improvement: Q1 at 14.2%, Q2 at 15.1%, '
        'Q3 at 16.3%, and Q4 at 17.8%. The full-year EBITDA margin of 15.9% represents a '
        'new company record and positions us favorably against industry benchmarks of 13.2%.'
    )

    # Section 3: Operational Metrics
    doc.add_heading('Operational Metrics and KPIs', level=1)
    doc.add_paragraph(
        'Operational performance continued to strengthen across all key metrics. '
        'The company achieved a 99.7% on-time delivery rate, up from 98.2% in FY2024. '
        'Quality incidents decreased by 34%, and average resolution time dropped to 2.3 hours.'
    )

    doc.add_heading('Customer Acquisition Metrics', level=2)
    doc.add_paragraph(
        'The company onboarded 1,247 new customers in FY2025, a 19% increase over the prior year. '
        'Customer acquisition cost (CAC) decreased to $3,420 from $4,180, reflecting improved '
        'marketing efficiency and stronger brand recognition.'
    )

    doc.add_heading('New Customer Onboarding Rate', level=3)
    doc.add_paragraph(
        'Average time-to-value for new customers improved from 45 days to 28 days following '
        'the implementation of our automated onboarding platform. Customer success team NPS '
        'scores averaged 72 for the year, up from 64 in FY2024.'
    )

    doc.add_heading('Employee Productivity Index', level=2)
    doc.add_paragraph(
        'Revenue per employee increased to $324K from $298K, a 8.7% improvement. '
        'The engineering team shipped 47 major features across 12 product releases. '
        'Average sprint velocity improved by 15% following the adoption of new development tools.'
    )

    doc.add_heading('Supply Chain Efficiency', level=2)
    doc.add_paragraph(
        'Inventory turnover ratio improved to 8.2x from 6.9x. Warehouse utilization '
        'reached 91%, and shipping costs per unit decreased by 11.3%. The implementation '
        'of predictive analytics reduced stockout incidents by 42%.'
    )

    # Section 4: Market Analysis
    doc.add_heading('Market Analysis and Competitive Landscape', level=1)
    doc.add_paragraph(
        'The total addressable market grew to $12.4B in 2025, up 9.2% from $11.4B in 2024. '
        'Our market share increased to 6.8% from 6.1%. Key competitive dynamics include '
        'consolidation among mid-tier players and increased investment in AI-driven solutions.'
    )

    doc.add_heading('Digital Transformation Progress', level=2)
    doc.add_paragraph(
        'Phase II of the digital transformation initiative was completed on schedule and under '
        'budget. Key deliverables included the migration of 94% of workloads to cloud infrastructure, '
        'deployment of ML-based demand forecasting, and launch of the customer self-service portal.'
    )

    doc.add_heading('Risk Assessment Summary', level=2)
    doc.add_paragraph(
        'The enterprise risk register identified 23 material risks, down from 31 in FY2024. '
        'Cybersecurity posture improved with zero data breaches and successful completion of '
        'SOC 2 Type II and ISO 27001 certifications.'
    )

    doc.add_heading('Regulatory Compliance Status', level=2)
    doc.add_paragraph(
        'All regulatory audits passed without material findings. GDPR compliance program '
        'matured to Level 4, and new data privacy frameworks were implemented for APAC markets. '
        'Environmental compliance metrics exceeded targets across all manufacturing facilities.'
    )

    doc.add_heading('Server Uptime and Performance', level=3)
    doc.add_paragraph(
        'Platform availability reached 99.97%, exceeding the 99.95% SLA target. Average API '
        'response time improved to 42ms from 67ms. The infrastructure team completed migration '
        'to Kubernetes, reducing deployment time from 4 hours to 12 minutes.'
    )

    # Section 5: Strategic Recommendations
    doc.add_heading('Strategic Recommendations and Outlook', level=1)
    doc.add_paragraph(
        'Based on the analysis presented in this report, the leadership team recommends '
        'five strategic priorities for FY2026: (1) accelerate APAC expansion, (2) invest '
        'in AI/ML product capabilities, (3) pursue two strategic acquisitions, (4) expand '
        'the partner ecosystem, and (5) launch the sustainability certification program.'
    )

    doc.add_heading('Technology Infrastructure Updates', level=2)
    doc.add_paragraph(
        'Planned investments of $28M in technology infrastructure include data center '
        'expansion, edge computing deployment, and advanced security operations center. '
        'Expected ROI of 340% over three years based on efficiency gains and risk reduction.'
    )

    doc.add_heading('Talent Management Overview', level=2)
    doc.add_paragraph(
        'Headcount is projected to grow from 2,612 to 3,100 employees. Key hiring priorities '
        'include data scientists (25 positions), solution architects (18 positions), and '
        'regional sales managers (12 positions). Employee retention target is 92%.'
    )

    doc.add_heading('Sustainability Initiatives', level=2)
    doc.add_paragraph(
        'The company commits to achieving carbon neutrality by 2028. FY2026 initiatives include '
        'solar panel installation at three facilities, electric vehicle fleet transition, and '
        'supply chain sustainability audits for all Tier 1 suppliers.'
    )

    doc.add_heading('Carbon Footprint Reduction', level=3)
    doc.add_paragraph(
        'Total carbon emissions decreased by 18% in FY2025 to 12,400 metric tons CO2e. '
        'Scope 1 emissions fell 22% due to facility upgrades, while Scope 2 decreased 15% '
        'through renewable energy procurement. Scope 3 reduction targets are being finalized.'
    )

    doc.add_heading('Diversity and Inclusion Metrics', level=3)
    doc.add_paragraph(
        'Female representation in leadership roles increased to 38% from 33%. The company '
        'achieved pay equity across all demographics following the third-party audit in Q2. '
        'Employee resource groups grew to 12, with 64% of staff participating in at least one.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
