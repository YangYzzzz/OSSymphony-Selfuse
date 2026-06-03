"""
Initial Setup: Writer document with body text using single line spacing and no paragraph spacing.
Task ID: writer_fs_022
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_022'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Configure the 'Text Body' style with single line spacing and no paragraph spacing
    text_body_style = doc.styles['Body Text']
    pf = text_body_style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0  # single line spacing
    text_body_style.font.name = 'Liberation Serif'
    text_body_style.font.size = Pt(12)

    # Add a title using Heading 1
    heading = doc.add_heading('Quarterly Performance Review — Q1 2025', level=1)

    # Add body text paragraphs using 'Body Text' style
    body_paragraphs = [
        (
            "The first quarter of 2025 has been a period of significant growth for Meridian "
            "Technologies. Our engineering division delivered three major product releases ahead "
            "of schedule, while the marketing team successfully launched campaigns across twelve "
            "new regional markets. Revenue increased by 18.4% compared to the same quarter last "
            "year, exceeding our internal forecast of 15%."
        ),
        (
            "Employee engagement scores remained strong at 82%, reflecting the positive impact "
            "of our flexible work policy introduced in November 2024. The new mentorship program, "
            "led by Director of People Operations Anika Patel, has matched 147 mentor-mentee pairs "
            "across all departments. Early feedback indicates improved cross-functional collaboration "
            "and knowledge sharing."
        ),
        (
            "On the financial side, operating expenses were held to $4.2 million against a budget "
            "of $4.5 million, resulting in a favorable variance of $300,000. Chief Financial Officer "
            "Marcus Rivera attributes this to renegotiated vendor contracts and a 12% reduction in "
            "cloud infrastructure costs following the migration to a hybrid hosting model."
        ),
        (
            "Customer satisfaction metrics showed notable improvement. The Net Promoter Score rose "
            "from 61 to 68, driven primarily by faster response times from the support team and the "
            "introduction of an AI-assisted troubleshooting portal. Support ticket resolution time "
            "decreased from an average of 4.3 hours to 2.8 hours."
        ),
        (
            "Looking ahead to Q2, the leadership team has identified three strategic priorities: "
            "expanding our enterprise client base in the Asia-Pacific region, completing the beta "
            "launch of the Meridian Analytics Platform, and finalizing the acquisition of DataStream "
            "Solutions. Each initiative has a dedicated project lead and a timeline approved by the "
            "board of directors."
        ),
        (
            "In summary, Q1 2025 demonstrates that Meridian Technologies is well-positioned for "
            "sustained growth. The combination of strong financial performance, high employee morale, "
            "and improving customer relationships provides a solid foundation for the ambitious goals "
            "set for the remainder of the fiscal year."
        ),
    ]

    for text in body_paragraphs:
        para = doc.add_paragraph(text, style='Body Text')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
