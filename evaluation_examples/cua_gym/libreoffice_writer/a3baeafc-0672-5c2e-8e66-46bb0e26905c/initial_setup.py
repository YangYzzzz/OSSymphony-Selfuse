"""
Initial Setup: Tab stop price list document (initial state with space-separated items)
Task ID: osworld_writer_tabstop_split_line_004
Domain: libreoffice_writer

Creates a price list document where items and prices are separated by spaces (no tab stops).
The agent's task is to replace spaces with tabs and set proper tab stops.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_split_line_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Spring Harvest Market — Price List")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)

    # Subtitle / intro line
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run("Fresh seasonal produce, artisan goods, and farm staples.")
    intro_run.font.size = Pt(11)
    intro_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    intro_para.paragraph_format.space_after = Pt(18)

    # Section header
    section_para = doc.add_paragraph()
    section_run = section_para.add_run("Produce & Pantry Items")
    section_run.bold = True
    section_run.font.size = Pt(12)
    section_para.paragraph_format.space_after = Pt(6)

    # Price list lines — item names and prices separated by spaces (NO tab stops)
    # Each line: "Item Name          $Price"
    items = [
        ("Organic Heirloom Tomatoes (1 lb)", "$3.75"),
        ("Wildflower Honey (12 oz jar)", "$8.50"),
        ("Stone-Ground Whole Wheat Flour (5 lb)", "$6.20"),
        ("Free-Range Brown Eggs (dozen)", "$5.90"),
        ("Cold-Press Apple Cider (half gallon)", "$7.40"),
    ]

    for item_name, price in items:
        para = doc.add_paragraph()
        # Use multiple spaces to visually separate item name and price
        run = para.add_run(f"{item_name}   {price}")
        run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(4)
        # NO tab stops added here — that's what the agent must do

    # Footer note
    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_run = note_para.add_run("* Prices subject to change. All produce locally sourced.")
    note_run.italic = True
    note_run.font.size = Pt(10)
    note_para.paragraph_format.space_before = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
