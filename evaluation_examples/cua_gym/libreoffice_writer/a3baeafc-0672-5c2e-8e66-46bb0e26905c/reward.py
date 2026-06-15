"""
Reward Script: Set up tab stops in price list paragraph
Task ID: osworld_writer_tabstop_split_line_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All 5 price-list paragraphs have a RIGHT tab stop at ~14 cm
  Component 2 (0.3 pts): All 5 price-list paragraphs use a tab character to separate item from price
  Component 3 (0.2 pts): Price values appear directly after a tab character (confirming tab-based layout, not spaces)
"""

import os

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Cm

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_split_line_004'

# Expected prices for each paragraph (para indices 3-7 in the document)
EXPECTED_PRICES = ['$3.75', '$8.50', '$6.20', '$5.90', '$7.40']
# Acceptable range for the 14 cm RIGHT tab stop position (in EMU, ±72000 = 0.2 cm tolerance)
RIGHT_TABSTOP_TARGET_EMU = int(Cm(14))   # 5040000
RIGHT_TABSTOP_TOLERANCE_EMU = 72000      # 0.2 cm tolerance

# Paragraph indices of the 5 price list items in the document
PRICE_PARA_INDICES = [3, 4, 5, 6, 7]


def has_right_tab_at_14cm(para):
    """Check if a paragraph has a RIGHT tab stop approximately at 14 cm."""
    for ts in para.paragraph_format.tab_stops:
        if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
            continue
        if ts.alignment == WD_TAB_ALIGNMENT.RIGHT:
            if abs(ts.position - RIGHT_TABSTOP_TARGET_EMU) <= RIGHT_TABSTOP_TOLERANCE_EMU:
                return True, ts.position
    return False, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Sanity check: document should have at least 8 paragraphs (indices 0-7)
    if len(doc.paragraphs) < 8:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 8")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: All 5 price-list paragraphs have a RIGHT tab stop at ~14 cm
    #              (0.5 points)
    # -----------------------------------------------------------------------
    try:
        paras_with_right_tab = 0
        right_tab_details = []
        for idx in PRICE_PARA_INDICES:
            para = doc.paragraphs[idx]
            found, pos = has_right_tab_at_14cm(para)
            if found:
                paras_with_right_tab += 1
                right_tab_details.append(f"para[{idx}]@{pos/360000:.2f}cm")
            else:
                right_tab_details.append(f"para[{idx}]:MISSING")

        if paras_with_right_tab == 5:
            total_score += 0.5
            print(f"PASS: Component 1 — All 5 price paragraphs have RIGHT tab stop at ~14 cm ({right_tab_details})")
        elif paras_with_right_tab > 0:
            partial = round(0.5 * paras_with_right_tab / 5, 2)
            total_score += partial
            print(f"PARTIAL: Component 1 — {paras_with_right_tab}/5 paragraphs have RIGHT tab at ~14 cm (+{partial}) ({right_tab_details})")
        else:
            print(f"FAIL: Component 1 — No price paragraph has RIGHT tab stop at ~14 cm ({right_tab_details})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: All 5 price-list paragraphs use a tab character (\t)
    #              between item name and price (not just spaces)
    #              (0.3 points)
    # -----------------------------------------------------------------------
    try:
        paras_with_tab_char = 0
        tab_char_details = []
        for idx in PRICE_PARA_INDICES:
            para = doc.paragraphs[idx]
            if '\t' in para.text:
                paras_with_tab_char += 1
                tab_char_details.append(f"para[{idx}]:OK")
            else:
                tab_char_details.append(f"para[{idx}]:NO_TAB_CHAR")

        if paras_with_tab_char == 5:
            total_score += 0.3
            print(f"PASS: Component 2 — All 5 price paragraphs use tab character separator ({tab_char_details})")
        elif paras_with_tab_char > 0:
            partial = round(0.3 * paras_with_tab_char / 5, 2)
            total_score += partial
            print(f"PARTIAL: Component 2 — {paras_with_tab_char}/5 paragraphs use tab char (+{partial}) ({tab_char_details})")
        else:
            print(f"FAIL: Component 2 — No price paragraph uses tab character separator ({tab_char_details})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Price values appear directly after a tab character (\t),
    #              confirming tab-based layout (not spaces before price).
    #              FAILS on initial (spaces used), PASSES on golden (tab before price).
    #              (0.2 points)
    # -----------------------------------------------------------------------
    try:
        prices_after_tab = 0
        price_tab_details = []
        for idx, expected_price in zip(PRICE_PARA_INDICES, EXPECTED_PRICES):
            para = doc.paragraphs[idx]
            # Check that the price appears right after a tab character
            if '\t' + expected_price in para.text:
                prices_after_tab += 1
                price_tab_details.append(f"para[{idx}]:OK")
            else:
                price_tab_details.append(f"para[{idx}]:FAIL")

        if prices_after_tab == 5:
            total_score += 0.2
            print(f"PASS: Component 3 — All 5 price values appear directly after tab char ({price_tab_details})")
        elif prices_after_tab > 0:
            partial = round(0.2 * prices_after_tab / 5, 2)
            total_score += partial
            print(f"PARTIAL: Component 3 — {prices_after_tab}/5 prices after tab char (+{partial}) ({price_tab_details})")
        else:
            print(f"FAIL: Component 3 — No price value found directly after tab character ({price_tab_details})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
