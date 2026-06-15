"""
Initial Setup: Technical Report with Appendix A (portrait throughout, no landscape)
Task ID: writer_page_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_073'
OUTPUT = f'{WORKDIR}/technical_report.docx'

# A4 dimensions in EMU (English Metric Units)
A4_WIDTH = Cm(21.0)
A4_HEIGHT = Cm(29.7)
MARGIN = Cm(2.54)


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_section_portrait(section):
    """Configure section as A4 portrait with 2.54cm margins."""
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN


def add_heading(doc, text, level):
    """Add a heading paragraph."""
    para = doc.add_heading(text, level=level)
    return para


def add_body_paragraph(doc, text):
    """Add a normal body paragraph."""
    para = doc.add_paragraph(text)
    return para


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Configure the single default section as A4 portrait ---
    section = doc.sections[0]
    set_section_portrait(section)

    # ===== COVER / TITLE =====
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run('Advanced Systems Engineering\nTechnical Report TR-2025-047')
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run2 = subtitle.add_run('Distributed Fault-Tolerant Architecture for Real-Time Data Processing')
    run2.italic = True
    run2.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.paragraph_format.space_after = Pt(6)
    meta.add_run('Prepared by: Dr. Elena Vasquez, Principal Engineer\nDate: March 2025\nVersion: 2.1 Final')

    doc.add_page_break()

    # ===== EXECUTIVE SUMMARY (page 2) =====
    add_heading(doc, 'Executive Summary', level=1)

    add_body_paragraph(doc,
        'This report presents the design, implementation, and validation findings of the Distributed '
        'Fault-Tolerant Architecture (DFTA) developed for high-throughput real-time data processing '
        'environments. The system was engineered to meet strict latency targets of under 50 milliseconds '
        'for 99.9% of transactions, while sustaining a throughput of 2.5 million events per second across '
        'a heterogeneous cluster of 48 processing nodes.')

    add_body_paragraph(doc,
        'Key achievements include: (1) a 34% reduction in end-to-end latency compared to the baseline '
        'monolithic architecture; (2) zero data loss over 720 hours of continuous stress testing; '
        '(3) automatic failover within 250 ms on node failure; and (4) linear horizontal scalability '
        'demonstrated up to 96 nodes.')

    add_body_paragraph(doc,
        'The DFTA leverages a novel consensus protocol — termed Adaptive Quorum Replication (AQR) — '
        'which dynamically adjusts replication factors based on network partition probability estimates '
        'derived from rolling latency histograms. Full technical details are presented in Sections 3 '
        'through 6.')

    doc.add_page_break()

    # ===== TABLE OF CONTENTS (page 3) =====
    add_heading(doc, 'Table of Contents', level=1)

    toc_entries = [
        ('1', 'Introduction', '4'),
        ('2', 'System Requirements', '5'),
        ('3', 'Architecture Overview', '6'),
        ('4', 'Component Design', '7'),
        ('5', 'Performance Evaluation', '8'),
        ('6', 'Conclusion', '9'),
        ('Appendix A', 'Raw Benchmark Data Tables', '10'),
    ]
    for num, title_text, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f'{num}  ')
        run_num.bold = True
        run_title = p.add_run(title_text)
        run_dots = p.add_run('.' * (50 - len(num) - len(title_text)))
        run_page = p.add_run(f'  {page}')

    doc.add_page_break()

    # ===== SECTION 1: INTRODUCTION (page 4) =====
    add_heading(doc, '1. Introduction', level=1)

    add_body_paragraph(doc,
        'Modern enterprise data pipelines must accommodate exponentially growing data volumes while '
        'maintaining strict service level agreements (SLAs). The financial services sector, smart '
        'manufacturing, and real-time analytics platforms all demand sub-100 ms processing latencies '
        'with five-nines availability (99.999%).')

    add_body_paragraph(doc,
        'Prior approaches based on centralised message brokers (e.g., Apache Kafka clusters running '
        'a single coordinator) suffer from well-documented scalability bottlenecks and split-brain '
        'vulnerabilities under network partitions. The DFTA addresses these shortcomings by eliminating '
        'the central coordinator in favour of a leaderless peer-to-peer topology with deterministic '
        'conflict resolution.')

    add_body_paragraph(doc,
        'The remainder of this report is organised as follows. Section 2 specifies the functional and '
        'non-functional requirements. Section 3 provides an architectural overview. Sections 4 and 5 '
        'detail the component design and performance evaluation respectively. Section 6 concludes with '
        'recommendations for production deployment. Appendix A contains the raw benchmark data tables.')

    doc.add_page_break()

    # ===== SECTION 2: SYSTEM REQUIREMENTS (page 5) =====
    add_heading(doc, '2. System Requirements', level=1)

    add_heading(doc, '2.1 Functional Requirements', level=2)

    add_body_paragraph(doc,
        'The system shall accept event streams encoded in Apache Avro format, provide at-least-once '
        'delivery guarantees with idempotent consumer support, persist all processed events in a '
        'distributed write-ahead log for 90 days, and expose a gRPC API for producer and consumer '
        'registration.')

    add_body_paragraph(doc,
        'The system shall support dynamic topic partitioning without service interruption, provide '
        'per-partition ordering guarantees, and publish Prometheus-compatible telemetry metrics at '
        '15-second intervals.')

    add_heading(doc, '2.2 Non-Functional Requirements', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers_row = table.rows[0]
    for i, h in enumerate(['Attribute', 'Target', 'Measurement Method']):
        cell = headers_row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    nfr_data = [
        ('Throughput', '2.5M events/sec', 'Apache JMeter load test'),
        ('Latency P99', '< 50 ms', 'Percentile histogram'),
        ('Availability', '99.999%', 'SLO tracking over 12 months'),
        ('RTO', '< 250 ms', 'Simulated node kill'),
        ('RPO', '0 events', 'Post-failover log comparison'),
    ]
    for r, (attr, target, method) in enumerate(nfr_data, 1):
        row = table.rows[r]
        row.cells[0].text = attr
        row.cells[1].text = target
        row.cells[2].text = method

    doc.add_page_break()

    # ===== SECTION 3: ARCHITECTURE OVERVIEW (page 6) =====
    add_heading(doc, '3. Architecture Overview', level=1)

    add_body_paragraph(doc,
        'The DFTA consists of four logical layers: (1) the Ingestion Layer, responsible for validating '
        'and serialising incoming events; (2) the Routing Layer, which assigns events to partitions '
        'using consistent hashing; (3) the Storage Layer, comprising the distributed write-ahead log; '
        'and (4) the Delivery Layer, which manages consumer group state and offset tracking.')

    add_body_paragraph(doc,
        'Each processing node participates equally in all four layers. Nodes form a ring topology '
        'maintained by a gossip protocol with O(log N) message complexity. Each event is replicated '
        'to three contiguous nodes on the ring before an acknowledgement is returned to the producer.')

    add_body_paragraph(doc,
        'The Adaptive Quorum Replication (AQR) protocol monitors rolling p99 latency for inter-node '
        'replication. When latency exceeds a configurable threshold (default 35 ms), AQR temporarily '
        'reduces the replication factor from three to two for the affected ring segment, prioritising '
        'availability over durability. The decision is recorded in the audit log and reversed '
        'automatically once latency stabilises below threshold for 60 consecutive seconds.')

    doc.add_page_break()

    # ===== SECTION 4: COMPONENT DESIGN (page 7) =====
    add_heading(doc, '4. Component Design', level=1)

    add_heading(doc, '4.1 Ingestion Gateway', level=2)

    add_body_paragraph(doc,
        'The Ingestion Gateway is a stateless service deployed as a Kubernetes DaemonSet. Each instance '
        'maintains a pool of 256 goroutines for concurrent request handling. Incoming gRPC requests are '
        'validated against the Avro schema registry and assigned a monotonically increasing sequence '
        'number within a 64-bit namespace partitioned by producer ID.')

    add_heading(doc, '4.2 Partition Router', level=2)

    add_body_paragraph(doc,
        'The Partition Router implements rendezvous (highest random weight) hashing for partition '
        'assignment. This algorithm provides superior load balance compared to modular hashing when '
        'the number of partitions changes, reducing data movement during cluster rebalancing to '
        'approximately 1/N of total data (where N is the new cluster size).')

    add_heading(doc, '4.3 Write-Ahead Log', level=2)

    add_body_paragraph(doc,
        'Each node maintains a local write-ahead log implemented as an append-only segmented file '
        'structure. Segments are 256 MB in size and compressed using LZ4 before being flushed to disk. '
        'The log index is maintained in memory using a skip-list for O(log N) seek operations. '
        'Compaction runs on a configurable schedule (default: every 6 hours) and removes duplicate '
        'events using SHA-256 hashed keys.')

    doc.add_page_break()

    # ===== SECTION 5: PERFORMANCE EVALUATION (page 8) =====
    add_heading(doc, '5. Performance Evaluation', level=1)

    add_heading(doc, '5.1 Test Environment', level=2)

    add_body_paragraph(doc,
        'All benchmarks were conducted on a private cloud cluster of 48 nodes. Each node is equipped '
        'with dual Intel Xeon Gold 6342 processors (24 cores each), 256 GB DDR4 ECC RAM, '
        'and four NVMe SSDs in RAID-0 configuration yielding 14 GB/s sequential write throughput. '
        'Nodes are interconnected via 25 Gbps Ethernet with a full-bisection-bandwidth leaf-spine fabric.')

    add_heading(doc, '5.2 Throughput Results', level=2)

    add_body_paragraph(doc,
        'Under the standard load profile (80% writes, 20% reads, 1 KB average event size), the 48-node '
        'cluster sustained 2.73 million events per second — 9.2% above the 2.5M target. Peak throughput '
        'of 3.1M events/sec was observed during 15-second burst windows before the ingestion gateways '
        'applied back-pressure.')

    add_heading(doc, '5.3 Latency Results', level=2)

    add_body_paragraph(doc,
        'End-to-end p50 latency was 18 ms; p99 latency was 43 ms, within the 50 ms SLA target. '
        'The p99.9 latency was 78 ms, attributable to occasional GC pauses in the Java-based '
        'consumer group coordinator service. Replacing the coordinator with a Rust implementation '
        'is planned for version 2.2.')

    doc.add_page_break()

    # ===== SECTION 6: CONCLUSION (page 9) =====
    add_heading(doc, '6. Conclusion', level=1)

    add_body_paragraph(doc,
        'The Distributed Fault-Tolerant Architecture has successfully demonstrated compliance with all '
        'specified functional and non-functional requirements. The Adaptive Quorum Replication protocol '
        'provides a pragmatic balance between consistency and availability under network stress conditions.')

    add_body_paragraph(doc,
        'Recommended next steps for production deployment: (1) replace the Java consumer group '
        'coordinator with the Rust implementation to reduce p99.9 latency; (2) implement automated '
        'capacity planning based on rolling throughput trends; (3) conduct a third-party security audit '
        'of the gRPC authentication layer before external exposure.')

    add_body_paragraph(doc,
        'The raw benchmark data supporting the performance claims in Section 5 are provided in '
        'Appendix A for independent verification.')

    doc.add_page_break()

    # ===== APPENDIX A (page 10) =====
    # NOTE: Appendix A is simply a new paragraph with Heading 1 style.
    # NO page break is added before the heading here — the task requires the agent to add it.
    # The doc.add_page_break() above accounts for the page break ending Section 6.
    # Appendix A flows naturally after it (no additional explicit page style change).
    appendix_heading = doc.add_heading('Appendix A', level=1)

    add_body_paragraph(doc, 'Raw Benchmark Data Tables')

    add_body_paragraph(doc,
        'The following tables contain the complete benchmark data sets collected during the performance '
        'evaluation phase. All measurements are recorded in milliseconds unless otherwise stated.')

    # Throughput table
    add_heading(doc, 'Table A.1: Throughput Measurements (events/second)', level=2)
    tbl1 = doc.add_table(rows=7, cols=4)
    tbl1.style = 'Table Grid'
    tbl1_headers = ['Node Count', 'P50 Throughput', 'P95 Throughput', 'Peak Throughput']
    for i, h in enumerate(tbl1_headers):
        cell = tbl1.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    tbl1_data = [
        ('12', '680,000', '710,000', '790,000'),
        ('24', '1,360,000', '1,400,000', '1,560,000'),
        ('36', '2,050,000', '2,110,000', '2,380,000'),
        ('48', '2,730,000', '2,810,000', '3,100,000'),
        ('60', '3,410,000', '3,500,000', '3,870,000'),
        ('72', '4,080,000', '4,190,000', '4,640,000'),
    ]
    for r, row_data in enumerate(tbl1_data, 1):
        for c, val in enumerate(row_data):
            tbl1.rows[r].cells[c].text = val

    # Latency table
    add_heading(doc, 'Table A.2: Latency Measurements (milliseconds)', level=2)
    tbl2 = doc.add_table(rows=7, cols=5)
    tbl2.style = 'Table Grid'
    tbl2_headers = ['Node Count', 'P50', 'P90', 'P99', 'P99.9']
    for i, h in enumerate(tbl2_headers):
        cell = tbl2.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    tbl2_data = [
        ('12', '22', '38', '61', '112'),
        ('24', '20', '34', '52', '95'),
        ('36', '19', '31', '47', '84'),
        ('48', '18', '29', '43', '78'),
        ('60', '17', '27', '40', '71'),
        ('72', '16', '26', '38', '68'),
    ]
    for r, row_data in enumerate(tbl2_data, 1):
        for c, val in enumerate(row_data):
            tbl2.rows[r].cells[c].text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
