"""
Reward Script: Create an alphabetical index at the end of the document
Task ID: writer_rd_031
Domain: libreoffice_writer
Scoring:
  Component 1: XE index entry fields for key terms (0.30 pts)
  Component 2: INDEX field present in document (0.20 pts)
  Component 3: Alphabetical Index heading exists near end (0.15 pts)
  Component 4: Index lists the expected terms (0.20 pts)
  Component 5: Index terms are alphabetically sorted (0.15 pts)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_031'

EXPECTED_TERMS = [
    'GDP', 'inflation', 'monetary policy', 'fiscal policy',
    'supply chain', 'demand curve', 'equilibrium',
    'trade deficit', 'interest rate', 'unemployment'
]


def persist_app_state(domain: str):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Parse XML namespace
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body_xml = doc.element.body

    # ---------------------------------------------------------------
    # Component 1: XE index entry fields for key terms (0.30 points)
    # These mark terms for inclusion in the index.
    # Initial env has 0 XE fields; golden has 10.
    # ---------------------------------------------------------------
    try:
        instr_fields = body_xml.findall('.//w:instrText', ns)
        xe_entries = []
        for field in instr_fields:
            if field.text and 'XE' in field.text:
                xe_entries.append(field.text.strip())

        # Extract the term names from XE fields (format: XE "term")
        found_terms = set()
        for xe in xe_entries:
            match = re.search(r'XE\s+"([^"]+)"', xe)
            if match:
                found_terms.add(match.group(1).lower())

        # Count how many of the 10 expected terms are marked
        matched_count = 0
        for term in EXPECTED_TERMS:
            if term.lower() in found_terms:
                matched_count += 1

        if matched_count >= 8:
            print(f"PASS: Component 1 — {matched_count}/10 expected terms have XE entries (0.30 pts)")
            total_score += 0.30
        elif matched_count >= 5:
            partial = round(0.30 * (matched_count / 10), 2)
            print(f"PARTIAL: Component 1 — {matched_count}/10 terms have XE entries ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — Only {matched_count}/10 expected terms have XE entries. Found: {found_terms}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: INDEX field present in document (0.20 points)
    # The INDEX field generates the actual index listing.
    # Initial env has no INDEX field; golden has one.
    # ---------------------------------------------------------------
    try:
        index_fields = []
        for field in instr_fields:
            if field.text and 'INDEX' in field.text:
                index_fields.append(field.text.strip())

        if len(index_fields) >= 1:
            print(f"PASS: Component 2 — INDEX field found: '{index_fields[0]}' (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — No INDEX field found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Alphabetical Index heading near end (0.15 points)
    # A heading paragraph with "Alphabetical Index" or "Index" should
    # exist in the latter part of the document.
    # Initial env has no such heading; golden does.
    # ---------------------------------------------------------------
    try:
        paras = doc.paragraphs
        num_paras = len(paras)
        # Search in the last 30% of the document
        search_start = max(0, int(num_paras * 0.7))
        comp3_score = 0.0
        for p in paras[search_start:]:
            style_name = p.style.name if p.style else ''
            if 'Heading' in style_name and 'index' in p.text.lower():
                print(f"PASS: Component 3 — Found heading '{p.text}' with style '{style_name}' (0.15 pts)")
                comp3_score = 0.15
                break

        if comp3_score == 0.0:
            # Also check for any paragraph near end with "alphabetical index" text
            for p in paras[search_start:]:
                if 'alphabetical index' in p.text.lower():
                    style_name = p.style.name if p.style else ''
                    print(f"PASS: Component 3 — Found 'Alphabetical Index' paragraph (style: {style_name}) (0.15 pts)")
                    comp3_score = 0.15
                    break

        if comp3_score == 0.0:
            print(f"FAIL: Component 3 — No 'Alphabetical Index' heading found in latter part of document")
        total_score += comp3_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Index lists the expected terms (0.20 points)
    # After the index heading, paragraphs should contain the key terms
    # with page numbers. Initial env has no such content.
    # ---------------------------------------------------------------
    try:
        paras = doc.paragraphs
        num_paras = len(paras)
        search_start = max(0, int(num_paras * 0.7))

        # Find the index heading position
        index_heading_idx = None
        for i in range(search_start, num_paras):
            text_lower = paras[i].text.lower()
            if 'alphabetical index' in text_lower or ('index' in text_lower and 'Heading' in (paras[i].style.name if paras[i].style else '')):
                index_heading_idx = i
                break

        if index_heading_idx is not None:
            # Collect text from paragraphs after the heading
            index_text_lines = []
            for p in paras[index_heading_idx + 1:]:
                if p.text.strip():
                    index_text_lines.append(p.text.strip().lower())

            # Check how many expected terms appear in the index listing
            terms_in_index = 0
            for term in EXPECTED_TERMS:
                term_lower = term.lower()
                if any(term_lower in line for line in index_text_lines):
                    terms_in_index += 1

            if terms_in_index >= 8:
                print(f"PASS: Component 4 — {terms_in_index}/10 terms listed in the index (0.20 pts)")
                total_score += 0.20
            elif terms_in_index >= 5:
                partial = round(0.20 * (terms_in_index / 10), 2)
                print(f"PARTIAL: Component 4 — {terms_in_index}/10 terms in index ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Only {terms_in_index}/10 terms found in index listing")
        else:
            print(f"FAIL: Component 4 — Cannot evaluate: no index heading found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---------------------------------------------------------------
    # Component 5: Index terms are alphabetically sorted (0.15 points)
    # The terms in the index should appear in alphabetical order.
    # This only passes if there is an index with terms.
    # ---------------------------------------------------------------
    try:
        if index_heading_idx is not None and len(index_text_lines) >= 5:
            # Extract the term names (before any tab/page number)
            term_names = []
            for line in index_text_lines:
                # Terms typically appear as "term\tpage_number" or "term page_number"
                parts = line.split('\t')
                term_name = parts[0].strip()
                if term_name:
                    term_names.append(term_name)

            if len(term_names) >= 5:
                # Check if sorted alphabetically
                is_sorted = term_names == sorted(term_names)
                if is_sorted:
                    print(f"PASS: Component 5 — Index terms are alphabetically sorted (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 5 — Index terms not sorted. Order: {term_names[:5]}...")
            else:
                print(f"FAIL: Component 5 — Too few terms ({len(term_names)}) to verify sorting")
        else:
            print(f"FAIL: Component 5 — Cannot evaluate sorting: no index heading or insufficient terms")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist unsaved edits if LibreOffice is open
persist_app_state("libreoffice_writer")

# Run verification
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
