"""
Initial Setup: Create a 25-page economics textbook chapter document
Task ID: writer_rd_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_styled_heading(doc, text, level=1):
    """Add a heading with specific styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body_paragraph(doc, text, first_indent=True):
    """Add a body paragraph with standard formatting."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title page
    for _ in range(6):
        doc.add_paragraph('')

    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Chapter 7')
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.name = 'Times New Roman'

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('Macroeconomic Principles and Policy')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.name = 'Times New Roman'

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_before = Pt(36)
    author_run = author_para.add_run('Dr. Elena Vasquez')
    author_run.font.size = Pt(14)
    author_run.font.name = 'Times New Roman'
    author_run.italic = True

    edition_para = doc.add_paragraph()
    edition_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    edition_run = edition_para.add_run('Foundations of Economic Theory, 4th Edition')
    edition_run.font.size = Pt(12)
    edition_run.font.name = 'Times New Roman'

    # Page break after title
    doc.add_page_break()

    # ===== Section 1: Introduction =====
    add_styled_heading(doc, '7.1 Introduction to Macroeconomics', level=1)

    add_body_paragraph(doc,
        'Macroeconomics is the branch of economics that studies the behavior and performance '
        'of an economy as a whole. Unlike microeconomics, which focuses on individual markets '
        'and agents, macroeconomics examines aggregate indicators such as GDP, unemployment '
        'rates, and national income to understand how the overall economy functions. This chapter '
        'provides a comprehensive overview of the key concepts that shape modern macroeconomic '
        'theory and practice.')

    add_body_paragraph(doc,
        'The study of macroeconomics gained prominence during the Great Depression of the 1930s, '
        'when the classical economic models failed to explain the prolonged period of high '
        'unemployment and economic stagnation. John Maynard Keynes revolutionized economic '
        'thinking by arguing that aggregate demand was the primary driving force in an economy, '
        'and that government intervention could be used to stabilize output over the business cycle.')

    add_body_paragraph(doc,
        'In modern economies, policymakers rely on a range of tools to manage economic performance. '
        'These include monetary policy, which involves controlling the money supply and interest '
        'rates, and fiscal policy, which involves government spending and taxation decisions. '
        'Understanding how these tools interact with market forces is essential for anyone seeking '
        'to grasp the complexities of contemporary economic systems.')

    add_body_paragraph(doc,
        'This chapter will explore each of these concepts in detail, examining how they relate '
        'to one another and how they influence the everyday lives of citizens, businesses, and '
        'governments around the world. We will begin with the most fundamental measure of '
        'economic activity: Gross Domestic Product.')

    doc.add_page_break()

    # ===== Section 2: GDP =====
    add_styled_heading(doc, '7.2 Gross Domestic Product (GDP)', level=1)

    add_body_paragraph(doc,
        'GDP is the total monetary value of all finished goods and services produced within '
        'a country\'s borders in a specific time period. It serves as a comprehensive scorecard '
        'of a given country\'s economic health. GDP can be measured using three approaches: '
        'the production approach, the income approach, and the expenditure approach.')

    add_body_paragraph(doc,
        'The expenditure approach calculates GDP as the sum of consumer spending (C), business '
        'investment (I), government spending (G), and net exports (NX). This can be expressed '
        'as the fundamental identity: GDP = C + I + G + (X - M), where X represents exports '
        'and M represents imports. This identity holds because every dollar spent on final '
        'goods and services must be received as income by someone in the economy.')

    add_body_paragraph(doc,
        'Real GDP adjusts for inflation by using constant base-year prices, making it possible '
        'to compare economic output across different time periods. Nominal GDP, by contrast, '
        'uses current prices and can be misleading when prices change significantly. The GDP '
        'deflator, calculated as (Nominal GDP / Real GDP) x 100, provides a measure of the '
        'overall price level in the economy.')

    add_body_paragraph(doc,
        'While GDP is the most widely used measure of economic activity, it has significant '
        'limitations. It does not account for the distribution of income, environmental '
        'degradation, or non-market activities such as household labor. Alternative measures '
        'such as the Human Development Index (HDI) and Genuine Progress Indicator (GPI) '
        'attempt to address these shortcomings.')

    add_body_paragraph(doc,
        'The growth rate of GDP is one of the most closely watched economic indicators. '
        'When real GDP increases, the economy is said to be expanding; when it decreases '
        'for two or more consecutive quarters, the economy is in a recession. Policymakers '
        'use GDP data to make decisions about fiscal policy and monetary policy interventions.')

    add_body_paragraph(doc,
        'Per capita GDP divides the total GDP by the population, providing a rough measure '
        'of average economic output per person. Countries with high per capita GDP generally '
        'enjoy higher standards of living, though this measure also fails to capture income '
        'inequality within a nation. Economists often use purchasing power parity (PPP) '
        'adjustments to make more meaningful cross-country comparisons of GDP.')

    doc.add_page_break()

    # ===== Section 3: Inflation =====
    add_styled_heading(doc, '7.3 Inflation and Price Stability', level=1)

    add_body_paragraph(doc,
        'Inflation is defined as a sustained increase in the general price level of goods '
        'and services in an economy over a period of time. When the general price level rises, '
        'each unit of currency buys fewer goods and services; consequently, inflation reflects '
        'a reduction in the purchasing power of money. The Consumer Price Index (CPI) is the '
        'most commonly used measure of inflation in most countries.')

    add_body_paragraph(doc,
        'There are several theories about the causes of inflation. Demand-pull inflation occurs '
        'when aggregate demand grows faster than aggregate supply, pulling prices upward. '
        'Cost-push inflation results from increases in the costs of production, such as rising '
        'wages or raw material prices, which are passed on to consumers. Built-in inflation '
        'is driven by adaptive expectations, where workers demand higher wages anticipating '
        'future price increases, creating a self-reinforcing cycle.')

    add_body_paragraph(doc,
        'Central banks typically target a moderate rate of inflation, usually around 2% per year, '
        'as a balance between economic growth and price stability. Deflation, or falling prices, '
        'can be even more damaging than moderate inflation because it increases the real burden '
        'of debt and can lead to a deflationary spiral where consumers delay purchases in '
        'anticipation of lower future prices.')

    add_body_paragraph(doc,
        'Hyperinflation, defined as monthly inflation rates exceeding 50%, represents an extreme '
        'failure of monetary management. Historical examples include Weimar Germany in the 1920s, '
        'Zimbabwe in the late 2000s, and Venezuela in the 2010s. In each case, excessive money '
        'printing by the government was the primary cause, leading to a collapse of confidence '
        'in the national currency.')

    add_body_paragraph(doc,
        'The relationship between inflation and unemployment is described by the Phillips Curve, '
        'which suggests an inverse relationship between the two variables in the short run. '
        'However, the stagflation experienced in the 1970s, when both inflation and unemployment '
        'were high simultaneously, challenged this relationship and led to the development of '
        'the expectations-augmented Phillips Curve.')

    add_body_paragraph(doc,
        'Measuring inflation accurately is essential for economic policy. The CPI tracks the '
        'cost of a fixed basket of goods and services over time, but it may overstate inflation '
        'due to substitution bias, quality improvements, and the introduction of new goods. '
        'The Producer Price Index (PPI) measures price changes from the perspective of sellers '
        'and often serves as a leading indicator of consumer inflation.')

    doc.add_page_break()

    # ===== Section 4: Monetary Policy =====
    add_styled_heading(doc, '7.4 Monetary Policy', level=1)

    add_body_paragraph(doc,
        'Monetary policy refers to the actions undertaken by a nation\'s central bank to control '
        'the money supply and achieve macroeconomic goals that promote sustainable economic growth. '
        'In the United States, the Federal Reserve (the Fed) is responsible for implementing '
        'monetary policy. The European Central Bank (ECB) serves a similar function for the '
        'Eurozone countries.')

    add_body_paragraph(doc,
        'The primary tools of monetary policy include open market operations, the discount rate, '
        'and reserve requirements. Open market operations involve the buying and selling of '
        'government securities in the open market. When the central bank purchases securities, '
        'it injects money into the banking system, lowering the interest rate and encouraging '
        'borrowing and investment.')

    add_body_paragraph(doc,
        'The federal funds rate, which is the interest rate at which banks lend reserves to '
        'each other overnight, serves as the benchmark for all other interest rates in the '
        'economy. When the central bank lowers the interest rate, it becomes cheaper for '
        'businesses and consumers to borrow money, stimulating economic activity. Conversely, '
        'raising the interest rate makes borrowing more expensive, which tends to cool down '
        'an overheating economy.')

    add_body_paragraph(doc,
        'Quantitative easing (QE) is an unconventional monetary policy tool used when the '
        'interest rate is already near zero and traditional tools are insufficient. Under QE, '
        'the central bank purchases longer-term securities, such as mortgage-backed securities '
        'and corporate bonds, to lower long-term interest rates and increase the money supply. '
        'The Fed implemented QE extensively during the 2008 financial crisis and again during '
        'the COVID-19 pandemic.')

    add_body_paragraph(doc,
        'The effectiveness of monetary policy depends on several factors, including the '
        'transmission mechanism (how changes in the money supply affect the real economy), '
        'the expectations of economic agents, and the current state of the economy. In a '
        'liquidity trap, where the interest rate is at or near zero, increasing the money '
        'supply may have little effect on economic activity because banks are unwilling to lend '
        'and consumers are unwilling to borrow.')

    add_body_paragraph(doc,
        'Forward guidance is another tool used by central banks to influence expectations '
        'about future monetary policy. By communicating their intentions regarding future '
        'interest rate paths, central banks can influence long-term interest rates and investment '
        'decisions today. The credibility of the central bank is crucial for the effectiveness '
        'of forward guidance, which is why central bank independence is considered essential '
        'for sound monetary policy management.')

    doc.add_page_break()

    # ===== Section 5: Fiscal Policy =====
    add_styled_heading(doc, '7.5 Fiscal Policy', level=1)

    add_body_paragraph(doc,
        'Fiscal policy involves the use of government spending and taxation to influence the '
        'economy. Unlike monetary policy, which is typically managed by an independent central '
        'bank, fiscal policy is determined by the legislative and executive branches of government. '
        'The two main instruments of fiscal policy are government expenditure and tax policy.')

    add_body_paragraph(doc,
        'Expansionary fiscal policy, which involves increasing government spending or reducing '
        'taxes, is used to stimulate economic growth during recessions. The logic is straightforward: '
        'when the government spends more or taxes less, aggregate demand increases, leading to '
        'higher output and employment. The multiplier effect amplifies the initial stimulus, as '
        'each dollar of government spending generates additional rounds of spending in the economy.')

    add_body_paragraph(doc,
        'Contractionary fiscal policy, which involves reducing government spending or increasing '
        'taxes, is used to cool down an overheating economy and control inflation. By reducing '
        'aggregate demand, the government can help prevent the economy from overheating and '
        'keep inflation in check. However, implementing contractionary fiscal policy is '
        'politically difficult, as spending cuts and tax increases are generally unpopular.')

    add_body_paragraph(doc,
        'The government budget deficit occurs when government spending exceeds tax revenue '
        'in a given fiscal year. Persistent deficits lead to the accumulation of national debt, '
        'which must be financed through the issuance of government bonds. The sustainability '
        'of fiscal policy depends on the relationship between the interest rate on government '
        'debt and the growth rate of GDP. If the growth rate exceeds the interest rate, the '
        'debt-to-GDP ratio can stabilize or even decline over time.')

    add_body_paragraph(doc,
        'Automatic stabilizers are fiscal mechanisms that automatically adjust government '
        'spending and tax revenue in response to economic conditions without requiring new '
        'legislation. Progressive income taxes, unemployment insurance, and social welfare '
        'programs all serve as automatic stabilizers. During recessions, tax revenue falls '
        'and welfare spending rises, providing a natural stimulus to the economy.')

    add_body_paragraph(doc,
        'The debate between proponents of active fiscal policy and those who favor a more '
        'hands-off approach remains central to macroeconomic discourse. Keynesian economists '
        'argue that fiscal policy is essential for managing aggregate demand, while supply-side '
        'economists emphasize the importance of tax incentives for promoting long-term growth '
        'through increased investment and productivity.')

    doc.add_page_break()

    # ===== Section 6: Supply and Demand =====
    add_styled_heading(doc, '7.6 Supply, Demand, and Market Equilibrium', level=1)

    add_body_paragraph(doc,
        'The concepts of supply and demand form the foundation of economic analysis. The '
        'demand curve shows the relationship between the price of a good and the quantity '
        'demanded by consumers, holding all other factors constant. The law of demand states '
        'that, ceteris paribus, as the price of a good increases, the quantity demanded '
        'decreases. This inverse relationship is represented by a downward-sloping demand curve.')

    add_body_paragraph(doc,
        'Several factors can shift the demand curve, including changes in consumer income, '
        'the prices of related goods (substitutes and complements), consumer preferences, '
        'population size, and expectations about future prices. A rightward shift indicates '
        'an increase in demand, while a leftward shift indicates a decrease in demand. It is '
        'important to distinguish between a movement along the demand curve (caused by a '
        'price change) and a shift of the entire curve (caused by non-price factors).')

    add_body_paragraph(doc,
        'The supply curve represents the relationship between the price of a good and the '
        'quantity that producers are willing to supply. According to the law of supply, as '
        'the price of a good increases, the quantity supplied also increases, resulting in '
        'an upward-sloping supply curve. Factors that can shift the supply curve include '
        'changes in input costs, technology, the number of sellers, government regulations, '
        'and expectations about future prices.')

    add_body_paragraph(doc,
        'Equilibrium occurs at the point where the supply curve and the demand curve intersect. '
        'At the equilibrium price, the quantity demanded by consumers equals the quantity supplied '
        'by producers, and there is no tendency for the price to change. If the market price is '
        'above equilibrium, a surplus exists, putting downward pressure on prices. If the market '
        'price is below equilibrium, a shortage exists, pushing prices upward.')

    add_body_paragraph(doc,
        'The concept of equilibrium is central to understanding how markets function. In '
        'competitive markets, the forces of supply and demand tend to push prices toward '
        'equilibrium. However, government interventions such as price ceilings and price '
        'floors can prevent markets from reaching equilibrium. Rent controls are a common '
        'example of a price ceiling, while minimum wage laws represent a price floor.')

    add_body_paragraph(doc,
        'Elasticity measures the responsiveness of one variable to changes in another. Price '
        'elasticity of demand measures how much the quantity demanded changes in response to '
        'a price change. Goods with elastic demand experience large quantity changes for small '
        'price changes, while goods with inelastic demand show relatively little change in '
        'quantity demanded even when prices change significantly. Understanding elasticity '
        'is crucial for businesses setting prices and for governments designing tax policy.')

    doc.add_page_break()

    # ===== Section 7: Supply Chain =====
    add_styled_heading(doc, '7.7 Global Supply Chains and Trade', level=1)

    add_body_paragraph(doc,
        'In today\'s interconnected global economy, the supply chain has become one of the '
        'most critical components of business operations. A supply chain encompasses all '
        'activities involved in the production and delivery of a product, from raw materials '
        'to the final consumer. Modern supply chains often span multiple countries and involve '
        'complex networks of suppliers, manufacturers, distributors, and retailers.')

    add_body_paragraph(doc,
        'The COVID-19 pandemic exposed the vulnerabilities of global supply chains, as '
        'lockdowns and border closures disrupted the flow of goods across borders. Many '
        'companies found that their just-in-time inventory systems, which had been designed '
        'to minimize costs, left them vulnerable to supply disruptions. The resulting shortages '
        'of everything from semiconductors to lumber highlighted the importance of supply chain '
        'resilience and risk management.')

    add_body_paragraph(doc,
        'International trade allows countries to specialize in producing goods and services '
        'in which they have a comparative advantage, leading to increased efficiency and '
        'higher overall output. The theory of comparative advantage, first articulated by '
        'David Ricardo in 1817, demonstrates that trade can be mutually beneficial even when '
        'one country is more efficient than another in producing all goods.')

    add_body_paragraph(doc,
        'A trade deficit occurs when a country imports more goods and services than it exports. '
        'The United States has run a persistent trade deficit since the 1970s, with the deficit '
        'reaching $948 billion in 2022. While a trade deficit is sometimes viewed negatively, '
        'economists debate whether it represents a genuine economic problem or simply reflects '
        'the attractiveness of a country as a destination for foreign investment.')

    add_body_paragraph(doc,
        'Tariffs, quotas, and other trade barriers are used by governments to protect domestic '
        'industries from foreign competition. While these measures may benefit specific sectors, '
        'they generally reduce overall economic efficiency by raising prices for consumers and '
        'distorting market signals. The World Trade Organization (WTO) works to reduce trade '
        'barriers and resolve trade disputes between nations.')

    add_body_paragraph(doc,
        'Supply chain management has become increasingly sophisticated with the adoption of '
        'technologies such as blockchain, artificial intelligence, and the Internet of Things. '
        'These technologies enable companies to track goods in real time, predict disruptions, '
        'and optimize logistics. However, the increasing complexity of global supply chains '
        'also creates new risks, including cybersecurity threats and geopolitical uncertainties '
        'that can affect the flow of goods and services.')

    doc.add_page_break()

    # ===== Section 8: Interest Rates =====
    add_styled_heading(doc, '7.8 Interest Rates and Financial Markets', level=1)

    add_body_paragraph(doc,
        'The interest rate is the cost of borrowing money or the return earned on savings and '
        'investments. It is one of the most important prices in the economy, influencing '
        'decisions about consumption, saving, and investment. Central banks set short-term '
        'interest rates as a primary tool of monetary policy, while long-term interest rates '
        'are determined by market forces, including expectations about future inflation and '
        'economic growth.')

    add_body_paragraph(doc,
        'The nominal interest rate is the stated rate on a loan or investment, while the real '
        'interest rate adjusts for inflation. The Fisher equation expresses this relationship: '
        'real interest rate = nominal interest rate - expected inflation rate. When inflation '
        'is high, the real interest rate may be negative, meaning that lenders lose purchasing '
        'power even though they receive positive nominal returns.')

    add_body_paragraph(doc,
        'The yield curve, which plots the relationship between interest rates and the maturity '
        'of government bonds, provides important information about market expectations. A normal '
        'yield curve slopes upward, indicating that investors demand higher interest rates for '
        'longer-term bonds to compensate for increased risk. An inverted yield curve, where '
        'short-term rates exceed long-term rates, has historically been a reliable predictor '
        'of economic recessions.')

    add_body_paragraph(doc,
        'Financial markets play a crucial role in the economy by channeling savings from '
        'households to productive investments. The stock market allows companies to raise '
        'capital by selling shares, while the bond market enables governments and corporations '
        'to borrow money. The efficiency of these markets depends on the flow of information '
        'and the transparency of pricing mechanisms.')

    add_body_paragraph(doc,
        'The relationship between interest rates and asset prices is inverse: when interest '
        'rates fall, bond prices rise, and stock prices generally increase as the present value '
        'of future earnings grows. This relationship explains why central bank interest rate '
        'decisions have such a significant impact on financial markets. Investors closely monitor '
        'Federal Reserve announcements for any indication of future changes in the interest rate '
        'path, often reacting within minutes of any policy statement.')

    add_body_paragraph(doc,
        'Credit markets, where borrowers and lenders interact, are essential for economic growth. '
        'When the interest rate is low, borrowing becomes more attractive, encouraging businesses '
        'to invest in new equipment, technology, and expansion. However, excessively low interest '
        'rates for extended periods can lead to asset bubbles, excessive risk-taking, and '
        'financial instability, as was observed in the lead-up to the 2008 financial crisis.')

    doc.add_page_break()

    # ===== Section 9: Unemployment =====
    add_styled_heading(doc, '7.9 Unemployment and Labor Markets', level=1)

    add_body_paragraph(doc,
        'Unemployment is one of the most closely monitored macroeconomic indicators, reflecting '
        'the health of the labor market and the overall economy. The unemployment rate is '
        'calculated as the percentage of the labor force that is actively seeking employment '
        'but unable to find work. A certain level of unemployment is considered natural and '
        'unavoidable in any dynamic economy.')

    add_body_paragraph(doc,
        'Economists distinguish between several types of unemployment. Frictional unemployment '
        'arises from the normal process of job searching and matching, as workers transition '
        'between jobs. Structural unemployment occurs when there is a mismatch between the '
        'skills workers possess and the skills employers need, often resulting from technological '
        'change or shifts in the structure of the economy.')

    add_body_paragraph(doc,
        'Cyclical unemployment is directly related to the business cycle. During recessions, '
        'as GDP contracts and firms reduce production, layoffs increase and the unemployment '
        'rate rises. During expansions, firms hire more workers and the unemployment rate falls. '
        'The natural rate of unemployment is the sum of frictional and structural unemployment '
        'and represents the lowest sustainable rate of unemployment in the economy.')

    add_body_paragraph(doc,
        'The labor force participation rate measures the percentage of the working-age population '
        'that is either employed or actively seeking employment. In recent decades, several '
        'trends have affected labor force participation, including the increasing participation '
        'of women, the retirement of the baby boom generation, and the growing number of '
        'discouraged workers who have stopped looking for employment.')

    add_body_paragraph(doc,
        'The relationship between unemployment and inflation, as described by the Phillips '
        'Curve, suggests that policymakers face a trade-off between the two. However, the '
        'long-run Phillips Curve is vertical at the natural rate of unemployment, implying '
        'that in the long run, monetary policy cannot permanently reduce unemployment below '
        'its natural rate without causing accelerating inflation.')

    add_body_paragraph(doc,
        'Government policies to reduce unemployment include fiscal stimulus, job training '
        'programs, unemployment insurance, and active labor market policies. The challenge '
        'for policymakers is to design interventions that address the root causes of '
        'unemployment without creating distortions in the labor market. In the context of '
        'rapid technological change and globalization, preparing workers for the jobs of the '
        'future has become an increasingly important policy priority.')

    doc.add_page_break()

    # ===== Section 10: Trade Deficit =====
    add_styled_heading(doc, '7.10 Trade Deficits and the Balance of Payments', level=1)

    add_body_paragraph(doc,
        'A trade deficit exists when a country\'s imports of goods and services exceed its '
        'exports. The trade balance is a key component of the current account, which also '
        'includes net income from abroad and net unilateral transfers. Understanding the trade '
        'deficit requires examining both the goods trade balance and the services trade balance, '
        'as many advanced economies run surpluses in services trade while maintaining deficits '
        'in goods trade.')

    add_body_paragraph(doc,
        'The balance of payments is a comprehensive record of all economic transactions between '
        'residents of a country and the rest of the world. It consists of the current account, '
        'the capital account, and the financial account. By definition, the balance of payments '
        'must sum to zero: a current account deficit must be offset by a financial account '
        'surplus, meaning that the country attracts net capital inflows from abroad.')

    add_body_paragraph(doc,
        'Several factors contribute to a trade deficit. A strong domestic currency makes imports '
        'cheaper and exports more expensive, widening the trade gap. High domestic demand, driven '
        'by robust consumer spending and investment, also tends to increase imports. Conversely, '
        'countries with high savings rates and export-oriented economies, such as China and '
        'Germany, tend to run trade surpluses.')

    add_body_paragraph(doc,
        'The relationship between the trade deficit and GDP is important for understanding the '
        'economic implications of trade imbalances. Net exports (exports minus imports) are a '
        'component of GDP, so a large trade deficit can reduce measured GDP growth. However, '
        'some economists argue that a trade deficit is not inherently harmful and may reflect '
        'strong investment opportunities that attract foreign capital.')

    add_body_paragraph(doc,
        'Exchange rates play a crucial role in determining trade flows. Under a floating exchange '
        'rate regime, the value of a currency is determined by supply and demand in foreign '
        'exchange markets. When a country runs a trade deficit, the supply of its currency in '
        'foreign exchange markets increases, which should in theory cause the currency to '
        'depreciate, making exports cheaper and imports more expensive, thereby correcting '
        'the trade imbalance over time.')

    add_body_paragraph(doc,
        'Trade policy debates often center on the question of whether trade deficits represent '
        'a threat to national economic security. Protectionists argue that persistent trade '
        'deficits lead to job losses in domestic manufacturing and increase dependence on foreign '
        'suppliers. Free trade advocates counter that trade deficits are the natural result of '
        'voluntary exchange and that restricting imports through tariffs and quotas reduces '
        'consumer welfare and economic efficiency.')

    doc.add_page_break()

    # ===== Section 11: Demand Curve Analysis =====
    add_styled_heading(doc, '7.11 Advanced Demand Analysis', level=1)

    add_body_paragraph(doc,
        'The demand curve is perhaps the most fundamental concept in economic analysis. At its '
        'core, the demand curve represents the willingness and ability of consumers to purchase '
        'a good at various price levels. The market demand curve is the horizontal summation of '
        'individual demand curves and reflects the total quantity demanded by all consumers at '
        'each price level.')

    add_body_paragraph(doc,
        'Consumer surplus, represented by the area above the price line and below the demand '
        'curve, measures the benefit consumers receive from purchasing a good at a price lower '
        'than their maximum willingness to pay. Changes in consumer surplus can be used to '
        'evaluate the welfare effects of government policies such as taxes, subsidies, and '
        'price controls.')

    add_body_paragraph(doc,
        'Income and substitution effects help explain the downward slope of the demand curve. '
        'When the price of a good falls, consumers can afford to buy more of it (income effect) '
        'and they tend to substitute it for relatively more expensive goods (substitution effect). '
        'For normal goods, both effects work in the same direction, reinforcing the law of demand. '
        'For inferior goods, the income effect works against the substitution effect, potentially '
        'weakening the price responsiveness of demand.')

    add_body_paragraph(doc,
        'The concept of the demand curve extends beyond individual goods to aggregate demand in '
        'the macroeconomy. The aggregate demand curve shows the total quantity of goods and '
        'services demanded at each overall price level. Shifts in aggregate demand, caused by '
        'changes in fiscal policy, monetary policy, consumer confidence, or global economic '
        'conditions, drive the business cycle and influence both GDP and inflation.')

    add_body_paragraph(doc,
        'Behavioral economics has challenged the traditional assumptions underlying the demand '
        'curve, showing that consumer decisions are often influenced by cognitive biases, social '
        'norms, and framing effects. Despite these critiques, the demand curve remains an '
        'indispensable tool for analyzing markets, predicting price changes, and evaluating '
        'policy interventions.')

    doc.add_page_break()

    # ===== Section 12: Equilibrium Models =====
    add_styled_heading(doc, '7.12 General Equilibrium and Market Dynamics', level=1)

    add_body_paragraph(doc,
        'The concept of equilibrium is central to economic theory. In a market context, '
        'equilibrium represents the state where the plans of buyers and sellers are consistent '
        'with each other, and there is no tendency for the market price to change. The notion '
        'of general equilibrium, developed by Leon Walras in the 19th century, extends this '
        'concept to consider the simultaneous equilibrium of all markets in an economy.')

    add_body_paragraph(doc,
        'Partial equilibrium analysis examines a single market in isolation, holding conditions '
        'in other markets constant. This approach is useful for understanding the effects of '
        'a specific policy change, such as a tax on cigarettes, on the market for that product. '
        'However, partial equilibrium analysis may miss important feedback effects that operate '
        'through other markets.')

    add_body_paragraph(doc,
        'In macroeconomics, the IS-LM model provides a framework for understanding how the '
        'goods market and the money market interact to determine equilibrium output and the '
        'interest rate simultaneously. The IS curve represents equilibrium in the goods market, '
        'where saving equals investment, while the LM curve represents equilibrium in the money '
        'market, where money demand equals money supply.')

    add_body_paragraph(doc,
        'The aggregate demand-aggregate supply (AD-AS) model is the primary tool for analyzing '
        'macroeconomic equilibrium. The intersection of the aggregate demand curve and the '
        'aggregate supply curve determines the equilibrium price level and the equilibrium level '
        'of real GDP. Short-run fluctuations in output and prices can be explained by shifts in '
        'aggregate demand or aggregate supply.')

    add_body_paragraph(doc,
        'Dynamic stochastic general equilibrium (DSGE) models represent the current frontier '
        'of macroeconomic modeling. These models incorporate rational expectations, market clearing, '
        'and stochastic shocks to analyze the effects of monetary policy and fiscal policy on '
        'the economy. While DSGE models are mathematically sophisticated, they have been '
        'criticized for their inability to predict financial crises and for their reliance on '
        'unrealistic assumptions about rational behavior.')

    add_body_paragraph(doc,
        'The search for economic equilibrium is not merely an academic exercise. Policymakers '
        'need to understand equilibrium dynamics to design effective interventions. Whether the '
        'economy is in equilibrium, moving toward equilibrium, or trapped in a suboptimal '
        'equilibrium has profound implications for decisions about monetary policy, fiscal '
        'policy, and structural reforms.')

    doc.add_page_break()

    # ===== Section 13: Policy Interactions =====
    add_styled_heading(doc, '7.13 Policy Coordination and Economic Stability', level=1)

    add_body_paragraph(doc,
        'Effective macroeconomic management requires careful coordination between monetary '
        'policy and fiscal policy. When these policies work in harmony, they can reinforce '
        'each other\'s effects, leading to more stable economic outcomes. For example, during '
        'a deep recession, expansionary monetary policy (lower interest rates) combined with '
        'expansionary fiscal policy (increased government spending) can provide a more powerful '
        'stimulus than either policy alone.')

    add_body_paragraph(doc,
        'However, policy coordination is often challenging in practice. Central banks and '
        'governments may have different objectives, time horizons, and political constraints. '
        'In many countries, central bank independence is designed to insulate monetary policy '
        'from political pressures, but this can sometimes lead to conflicts between fiscal and '
        'monetary authorities. The European sovereign debt crisis highlighted the difficulties '
        'of coordinating fiscal policy across multiple countries sharing a common monetary policy.')

    add_body_paragraph(doc,
        'The global financial crisis of 2008 demonstrated the importance of policy coordination '
        'on an international scale. As the crisis spread across borders, countries coordinated '
        'their responses through institutions like the G20 and the International Monetary Fund '
        '(IMF). Central banks around the world simultaneously lowered interest rates and '
        'implemented quantitative easing programs, while governments enacted fiscal stimulus '
        'packages to support aggregate demand and prevent a global depression.')

    add_body_paragraph(doc,
        'The interaction between monetary policy and the supply chain has become increasingly '
        'important in recent years. Supply chain disruptions can cause cost-push inflation that '
        'is difficult for monetary policy to address without causing a recession. The post-pandemic '
        'inflation surge illustrates this challenge: while some of the inflation was driven by '
        'demand-side factors that could be addressed by raising interest rates, much of it '
        'resulted from supply chain bottlenecks that were beyond the reach of monetary policy.')

    add_body_paragraph(doc,
        'Looking ahead, several structural changes in the global economy will shape the '
        'effectiveness of macroeconomic policy. Climate change, demographic shifts, technological '
        'disruption, and rising inequality all pose challenges that cannot be fully addressed '
        'by traditional monetary and fiscal policy tools. The relationship between GDP growth '
        'and human welfare is being reexamined, and new approaches to economic management that '
        'go beyond conventional measures of inflation and unemployment are being developed.')

    add_body_paragraph(doc,
        'The ongoing debates about the appropriate role of government in the economy reflect '
        'fundamental disagreements about the nature of equilibrium, the causes of market '
        'failures, and the effectiveness of policy interventions. Understanding these debates '
        'requires a solid grounding in the concepts covered in this chapter, from GDP measurement '
        'and inflation dynamics to the mechanics of monetary policy and fiscal policy, and from '
        'the behavior of the demand curve and supply chain to the determinants of the trade '
        'deficit, interest rate, and unemployment.')

    doc.add_page_break()

    # ===== Section 14: Case Studies =====
    add_styled_heading(doc, '7.14 Case Studies in Macroeconomic Policy', level=1)

    add_styled_heading(doc, '7.14.1 The Great Recession (2007-2009)', level=2)

    add_body_paragraph(doc,
        'The Great Recession was triggered by the collapse of the U.S. housing bubble and the '
        'subsequent failure of major financial institutions. The crisis exposed the fragility '
        'of the global financial system and the interconnectedness of modern supply chains. '
        'As credit markets froze, businesses could not obtain financing, leading to widespread '
        'layoffs and a sharp increase in unemployment. GDP contracted by 4.3% in the United '
        'States and by even more in some European countries.')

    add_body_paragraph(doc,
        'The policy response to the Great Recession involved unprecedented coordination between '
        'monetary policy and fiscal policy. The Federal Reserve cut the interest rate to near '
        'zero and implemented multiple rounds of quantitative easing. The U.S. government passed '
        'the American Recovery and Reinvestment Act of 2009, a $787 billion fiscal stimulus '
        'package designed to boost aggregate demand and prevent further job losses.')

    add_styled_heading(doc, '7.14.2 Japan\'s Lost Decades', level=2)

    add_body_paragraph(doc,
        'Japan\'s experience since the early 1990s illustrates the challenges of escaping a '
        'deflationary trap. After the collapse of its asset price bubble, Japan experienced '
        'prolonged stagnation characterized by low GDP growth, falling prices, and persistent '
        'unemployment. Despite extensive monetary policy easing, including zero interest rates '
        'and massive quantitative easing, Japan struggled to generate sustained inflation and '
        'economic growth for over two decades.')

    add_body_paragraph(doc,
        'Japan\'s experience highlighted the limitations of monetary policy when the economy is '
        'in a liquidity trap and consumer expectations are anchored at deflation. It also '
        'demonstrated the importance of fiscal policy in supporting demand, though Japan\'s '
        'massive fiscal stimulus programs contributed to a public debt exceeding 250% of GDP, '
        'raising concerns about long-term fiscal sustainability.')

    add_styled_heading(doc, '7.14.3 The European Sovereign Debt Crisis', level=2)

    add_body_paragraph(doc,
        'The European sovereign debt crisis, which began in 2010, exposed the tensions inherent '
        'in a monetary union without a fiscal union. Countries like Greece, Ireland, Portugal, '
        'and Spain experienced severe recessions, rising unemployment, and unsustainable '
        'government debt levels. The European Central Bank\'s monetary policy was constrained '
        'by the need to serve the entire Eurozone, while individual countries could not devalue '
        'their currencies to boost exports and reduce their trade deficits.')

    add_body_paragraph(doc,
        'The crisis was eventually contained through a combination of bailout programs, '
        'structural reforms, and ECB interventions, including the promise to do "whatever it '
        'takes" to preserve the euro. The experience highlighted the importance of fiscal '
        'discipline, the challenges of operating with a shared monetary policy but independent '
        'fiscal policies, and the complex relationship between interest rates, government debt, '
        'and economic growth.')

    add_styled_heading(doc, '7.14.4 Post-Pandemic Inflation (2021-2023)', level=2)

    add_body_paragraph(doc,
        'The economic recovery from the COVID-19 pandemic was accompanied by a sharp increase '
        'in inflation across major economies. In the United States, the annual inflation rate '
        'reached 9.1% in June 2022, the highest level in four decades. The causes of this '
        'inflation were multifaceted, including supply chain disruptions, expansionary fiscal '
        'policy (pandemic stimulus checks), accommodative monetary policy (near-zero interest '
        'rates and quantitative easing), and a rapid recovery in demand.')

    add_body_paragraph(doc,
        'Central banks responded by raising interest rates aggressively. The Federal Reserve '
        'raised the federal funds rate from near zero to over 5% in just 16 months, the fastest '
        'tightening cycle in decades. While higher interest rates helped cool inflation, they '
        'also raised concerns about a potential recession, increased the cost of servicing '
        'government debt, and exposed vulnerabilities in the banking sector, as evidenced by '
        'the failures of Silicon Valley Bank and other institutions in 2023.')

    doc.add_page_break()

    # ===== Section 15: Summary =====
    add_styled_heading(doc, '7.15 Chapter Summary', level=1)

    add_body_paragraph(doc,
        'This chapter has provided a comprehensive overview of the key concepts in macroeconomics. '
        'We began with GDP, the most widely used measure of economic activity, and examined its '
        'components, measurement methods, and limitations. We then explored inflation, discussing '
        'its causes, consequences, and the challenges it poses for monetary policy.')

    add_body_paragraph(doc,
        'The chapter examined the two main policy levers available to governments: monetary '
        'policy, implemented by central banks through control of the interest rate and money '
        'supply, and fiscal policy, implemented through government spending and taxation. We '
        'discussed how these policies interact and the challenges of policy coordination in '
        'a globalized economy.')

    add_body_paragraph(doc,
        'We analyzed the fundamental concepts of supply and demand, including the demand curve, '
        'the supply curve, and market equilibrium. These concepts were extended to the '
        'macroeconomic level through the aggregate demand-aggregate supply model, which provides '
        'a framework for understanding fluctuations in output and prices.')

    add_body_paragraph(doc,
        'The chapter also examined the roles of the supply chain in modern economies, the '
        'determinants and implications of the trade deficit, the importance of interest rates '
        'in financial markets, and the causes and consequences of unemployment. Each of these '
        'topics illustrates the interconnectedness of the macroeconomy and the complexity of '
        'economic policymaking.')

    add_body_paragraph(doc,
        'Finally, our case studies demonstrated how these concepts play out in practice, from '
        'the Great Recession and Japan\'s lost decades to the European sovereign debt crisis '
        'and the post-pandemic inflation surge. These episodes illustrate both the power and '
        'the limitations of macroeconomic policy in addressing real-world economic challenges.')

    # ===== Key Terms Box =====
    add_styled_heading(doc, 'Key Terms', level=2)

    terms = [
        ('GDP (Gross Domestic Product)', 'The total monetary value of all finished goods and services produced within a country\'s borders in a specific time period.'),
        ('Inflation', 'A sustained increase in the general price level of goods and services in an economy over time.'),
        ('Monetary Policy', 'Actions undertaken by a central bank to control the money supply and achieve macroeconomic goals.'),
        ('Fiscal Policy', 'The use of government spending and taxation to influence the economy.'),
        ('Supply Chain', 'The network of all activities involved in the production and delivery of a product, from raw materials to the final consumer.'),
        ('Demand Curve', 'A graph showing the relationship between the price of a good and the quantity demanded by consumers.'),
        ('Equilibrium', 'The state where the quantity demanded equals the quantity supplied, and there is no tendency for the price to change.'),
        ('Trade Deficit', 'The amount by which the cost of a country\'s imports exceeds the value of its exports.'),
        ('Interest Rate', 'The cost of borrowing money or the return earned on savings and investments.'),
        ('Unemployment', 'The state of being without a job while actively seeking employment; measured as a percentage of the labor force.'),
    ]

    for term_name, term_def in terms:
        p = doc.add_paragraph()
        run_term = p.add_run(term_name)
        run_term.bold = True
        run_term.font.name = 'Times New Roman'
        run_term.font.size = Pt(11)
        run_def = p.add_run(f' - {term_def}')
        run_def.font.name = 'Times New Roman'
        run_def.font.size = Pt(11)

    # ===== Review Questions =====
    doc.add_page_break()
    add_styled_heading(doc, 'Review Questions', level=2)

    questions = [
        'Explain the three approaches to measuring GDP and discuss the advantages and limitations of each.',
        'Compare and contrast demand-pull inflation and cost-push inflation. Provide examples of each.',
        'Describe the main tools of monetary policy and explain how they affect the interest rate and economic activity.',
        'Discuss the role of automatic stabilizers in fiscal policy. How do they differ from discretionary fiscal policy?',
        'Explain how disruptions in the supply chain can affect both GDP and inflation.',
        'Using the concept of the demand curve, explain how a change in consumer income affects the market for normal goods versus inferior goods.',
        'Define equilibrium in the context of market economics. What happens when the market is not in equilibrium?',
        'Analyze the causes and consequences of a persistent trade deficit for a country\'s economy.',
        'Explain the Fisher equation and discuss the difference between nominal and real interest rates.',
        'Distinguish between frictional, structural, and cyclical unemployment. Which type is most responsive to monetary policy?',
    ]

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f'{i}. ')
        run_num.bold = True
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(11)
        run_q = p.add_run(q)
        run_q.font.name = 'Times New Roman'
        run_q.font.size = Pt(11)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
