"""
Reward Script: Classic Italian Pasta Recipes booklet page
Task ID: writer_wf_021
Domain: libreoffice_writer
Scoring:
  Component 1: Two recipes with Heading 2 names (0.20)
  Component 2: Italic prep/cook time lines for each recipe (0.15)
  Component 3: Ingredients section with bulleted list (6+ items) per recipe (0.20)
  Component 4: Directions section with numbered list (5+ steps) per recipe (0.20)
  Component 5: Page break between recipes (0.10)
  Component 6: Header with booklet title on each page (0.15)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_021'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather structural information
    paragraphs = doc.paragraphs
    heading2_indices = []
    heading3_indices = []
    for i, p in enumerate(paragraphs):
        if p.style and p.style.name == 'Heading 2':
            heading2_indices.append(i)
        if p.style and p.style.name == 'Heading 3':
            heading3_indices.append(i)

    # =====================================================================
    # Component 1: Two recipes with Heading 2 names (0.20 points)
    # Each recipe should have a name as Heading 2. Need exactly 2 H2 headings
    # with non-empty text.
    # =====================================================================
    try:
        h2_with_text = [i for i in heading2_indices if paragraphs[i].text.strip()]
        if len(h2_with_text) >= 2:
            print(f"PASS: Component 1 — Found {len(h2_with_text)} Heading 2 recipes: "
                  f"{[paragraphs[i].text.strip()[:40] for i in h2_with_text[:2]]} (0.20 pts)")
            total_score += 0.20
        elif len(h2_with_text) == 1:
            print(f"FAIL: Component 1 — Only 1 Heading 2 found, need 2")
        else:
            print(f"FAIL: Component 1 — No Heading 2 paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # If we don't have at least 2 H2 headings, the rest of the scoring is meaningless
    if len(heading2_indices) < 2:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Build per-recipe paragraph ranges
    recipe_ranges = []
    for idx, h2_idx in enumerate(heading2_indices[:2]):
        end_idx = heading2_indices[idx + 1] if idx + 1 < len(heading2_indices) else len(paragraphs)
        recipe_ranges.append((h2_idx, end_idx))

    # =====================================================================
    # Component 2: Italic prep/cook time for each recipe (0.15 points)
    # The paragraph after each Heading 2 should contain italic text with
    # time information (prep time, cook time).
    # =====================================================================
    try:
        italic_recipes = 0
        for r_idx, (start, end) in enumerate(recipe_ranges):
            # Look for an italic paragraph near the start (within first 3 paragraphs after H2)
            found_italic_time = False
            for pi in range(start + 1, min(start + 4, end)):
                p = paragraphs[pi]
                text_lower = p.text.lower()
                has_time_info = ('prep' in text_lower or 'cook' in text_lower or 'time' in text_lower or 'minutes' in text_lower)
                has_italic = any(r.italic for r in p.runs if r.text.strip())
                if has_time_info and has_italic:
                    found_italic_time = True
                    break
            if found_italic_time:
                italic_recipes += 1

        if italic_recipes >= 2:
            print(f"PASS: Component 2 — Both recipes have italic prep/cook time info (0.15 pts)")
            total_score += 0.15
        elif italic_recipes == 1:
            print(f"PARTIAL: Component 2 — Only 1 of 2 recipes has italic time info (0.075 pts)")
            total_score += 0.075
        else:
            print(f"FAIL: Component 2 — No recipes have italic prep/cook time info")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =====================================================================
    # Component 3: Ingredients section with 6+ bulleted items per recipe (0.20 points)
    # Each recipe should have an "Ingredients" heading (H3) followed by List Bullet items.
    # =====================================================================
    try:
        recipes_with_ingredients = 0
        for r_idx, (start, end) in enumerate(recipe_ranges):
            # Find "Ingredients" heading within this recipe range
            ingredients_idx = None
            for pi in range(start, end):
                p = paragraphs[pi]
                if p.text.strip().lower() == 'ingredients' and p.style and 'Heading' in p.style.name:
                    ingredients_idx = pi
                    break

            if ingredients_idx is None:
                print(f"  Recipe {r_idx+1}: No 'Ingredients' heading found")
                continue

            # Count bullet items after Ingredients heading until next heading or end
            bullet_count = 0
            for pi in range(ingredients_idx + 1, end):
                p = paragraphs[pi]
                if p.style and 'Heading' in p.style.name:
                    break
                if p.style and p.style.name == 'List Bullet' and p.text.strip():
                    bullet_count += 1

            if bullet_count >= 6:
                recipes_with_ingredients += 1
                print(f"  Recipe {r_idx+1}: {bullet_count} bullet ingredients (>= 6)")
            else:
                print(f"  Recipe {r_idx+1}: Only {bullet_count} bullet ingredients (need >= 6)")

        if recipes_with_ingredients >= 2:
            print(f"PASS: Component 3 — Both recipes have 6+ bulleted ingredients (0.20 pts)")
            total_score += 0.20
        elif recipes_with_ingredients == 1:
            print(f"PARTIAL: Component 3 — Only 1 of 2 recipes has 6+ ingredients (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No recipe has 6+ bulleted ingredients")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =====================================================================
    # Component 4: Directions section with 5+ numbered steps per recipe (0.20 points)
    # Each recipe should have a "Directions" heading (H3) followed by List Number items.
    # =====================================================================
    try:
        recipes_with_directions = 0
        for r_idx, (start, end) in enumerate(recipe_ranges):
            # Find "Directions" heading within this recipe range
            directions_idx = None
            for pi in range(start, end):
                p = paragraphs[pi]
                if p.text.strip().lower() == 'directions' and p.style and 'Heading' in p.style.name:
                    directions_idx = pi
                    break

            if directions_idx is None:
                print(f"  Recipe {r_idx+1}: No 'Directions' heading found")
                continue

            # Count numbered items after Directions heading until next heading or end
            number_count = 0
            for pi in range(directions_idx + 1, end):
                p = paragraphs[pi]
                if p.style and 'Heading' in p.style.name:
                    break
                if p.style and p.style.name == 'List Number' and p.text.strip():
                    number_count += 1

            if number_count >= 5:
                recipes_with_directions += 1
                print(f"  Recipe {r_idx+1}: {number_count} numbered directions (>= 5)")
            else:
                print(f"  Recipe {r_idx+1}: Only {number_count} numbered directions (need >= 5)")

        if recipes_with_directions >= 2:
            print(f"PASS: Component 4 — Both recipes have 5+ numbered directions (0.20 pts)")
            total_score += 0.20
        elif recipes_with_directions == 1:
            print(f"PARTIAL: Component 4 — Only 1 of 2 recipes has 5+ directions (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — No recipe has 5+ numbered directions")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =====================================================================
    # Component 5: Page break between recipes (0.10 points)
    # There should be at least one page break between the first and second recipe.
    # =====================================================================
    try:
        page_break_found = False
        first_h2 = heading2_indices[0]
        second_h2 = heading2_indices[1]

        # Check for page breaks (both br type=page and page_break_before) between recipes
        for pi in range(first_h2, second_h2 + 1):
            p = paragraphs[pi]
            # Check page_break_before
            if p.paragraph_format.page_break_before:
                page_break_found = True
                break
            # Check for w:br type=page in runs
            for run in p.runs:
                for br in run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                    btype = br.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', '')
                    if btype == 'page':
                        page_break_found = True
                        break
                if page_break_found:
                    break
            if page_break_found:
                break

        if page_break_found:
            print(f"PASS: Component 5 — Page break found between recipes (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — No page break found between the two recipes")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # =====================================================================
    # Component 6: Header with booklet title (0.15 points)
    # The document header should contain "Classic Italian Pasta Recipes" or similar.
    # =====================================================================
    try:
        header_found = False
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                header_text = ' '.join(p.text for p in section.header.paragraphs).strip()
                if header_text:
                    # Check that the header contains key words from the booklet title
                    ht_lower = header_text.lower()
                    has_italian = 'italian' in ht_lower
                    has_pasta = 'pasta' in ht_lower
                    has_recipe = 'recipe' in ht_lower
                    if has_italian and has_pasta and has_recipe:
                        header_found = True
                        print(f"  Header text: {header_text!r}")
                        break

        if header_found:
            print(f"PASS: Component 6 — Header contains booklet title (0.15 pts)")
            total_score += 0.15
        else:
            # Collect all header texts for debugging
            all_headers = []
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    all_headers.append(' '.join(p.text for p in section.header.paragraphs).strip())
            print(f"FAIL: Component 6 — Header missing booklet title. Found headers: {all_headers}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
