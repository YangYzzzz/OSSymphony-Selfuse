"""
Reward Script: Table chapter-based numbering in Writer document
Task ID: writer_acad_056
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Table caption paragraphs use chapter-based numbering (Table X.Y format)
  Component 2 (0.3): Correct chapter assignment and numbering reset per chapter
  Component 3 (0.2): In-text cross-references updated to chapter-based numbers
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_056'

# Expected chapter-based caption mapping (paragraph index -> expected label)
# Chapter 1 has 2 tables: Table 1.1, Table 1.2
# Chapter 2 has 2 tables: Table 2.1, Table 2.2
# Chapter 3 has 2 tables: Table 3.1, Table 3.2
EXPECTED_CAPTIONS = {
    'Table 1.1': 'Summary of Recent Traffic Prediction Studies',
    'Table 1.2': 'Sensor Deployment Summary Across Study Areas',
    'Table 2.1': 'Hyperparameter Configurations for Each Model',
    'Table 2.2': 'Feature Engineering Summary',
    'Table 3.1': 'Prediction Performance Comparison (15-min Horizon)',
    'Table 3.2': 'Statistical Significance Tests (Paired t-test, p-values)',
}

# Expected in-text references (non-caption paragraphs that mention tables)
EXPECTED_REFS = {
    'Table 1.1': True,  # referenced in body text
    'Table 1.2': False, # only appears as caption
    'Table 2.1': False,
    'Table 2.2': False,
    'Table 3.1': False,
    'Table 3.2': True,  # referenced in body text paragraph 26
}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Identify caption paragraphs: those starting with "Table " followed by a number pattern
    caption_pattern = re.compile(r'^Table\s+(\d+(?:\.\d+)?)\s*:')
    chapter_pattern = re.compile(r'^Table\s+(\d+)\.(\d+)\s*:')

    captions_found = []
    for para in doc.paragraphs:
        m = caption_pattern.match(para.text)
        if m:
            captions_found.append(para.text)

    print(f"INFO: Found {len(captions_found)} table captions")
    for c in captions_found:
        print(f"  Caption: {c[:80]}")

    # Component 1: Table captions use chapter-based numbering format (0.5 points)
    # Check that all 6 captions use the "Table X.Y" format (with a dot)
    try:
        chapter_based_count = 0
        for cap_text in captions_found:
            m = chapter_pattern.match(cap_text)
            if m:
                chapter_based_count += 1

        if chapter_based_count >= 6:
            print(f"PASS: Component 1 — All {chapter_based_count} captions use chapter-based numbering (0.5 pts)")
            total_score += 0.5
        elif chapter_based_count >= 3:
            partial = round(0.5 * (chapter_based_count / 6), 2)
            print(f"PARTIAL: Component 1 — {chapter_based_count}/6 captions use chapter-based format ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — Only {chapter_based_count}/6 captions use chapter-based numbering")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Correct chapter assignment and numbering reset (0.3 points)
    # Verify: Ch1 -> 1.1, 1.2; Ch2 -> 2.1, 2.2; Ch3 -> 3.1, 3.2
    try:
        # Build ordered list of (chapter_num, table_num) from captions
        found_labels = []
        for cap_text in captions_found:
            m = chapter_pattern.match(cap_text)
            if m:
                ch = int(m.group(1))
                tb = int(m.group(2))
                found_labels.append((ch, tb))

        expected_labels = [(1, 1), (1, 2), (2, 1), (2, 2), (3, 1), (3, 2)]

        if found_labels == expected_labels:
            print(f"PASS: Component 2 — Chapter assignment and numbering reset correct (0.3 pts)")
            total_score += 0.3
        else:
            # Partial credit: count how many match in sequence
            matches = sum(1 for a, b in zip(found_labels, expected_labels) if a == b)
            if matches >= 3:
                partial = round(0.3 * (matches / 6), 2)
                print(f"PARTIAL: Component 2 — {matches}/6 labels match expected sequence ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — Expected {expected_labels}, found {found_labels}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: In-text cross-references updated (0.2 points)
    # Check that body text (non-caption) paragraphs reference "Table X.Y" not "Table N"
    try:
        # Collect all table references from non-caption paragraphs
        ref_pattern = re.compile(r'Table\s+(\d+(?:\.\d+)?)')
        sequential_ref_pattern = re.compile(r'Table\s+(\d+)(?!\.\d)')

        non_caption_refs_chapter = 0
        non_caption_refs_sequential = 0

        for para in doc.paragraphs:
            # Skip caption paragraphs
            if caption_pattern.match(para.text):
                continue
            # Find all references
            chapter_refs = re.findall(r'Table\s+\d+\.\d+', para.text)
            seq_refs = re.findall(r'Table\s+\d+(?!\.\d)', para.text)

            non_caption_refs_chapter += len(chapter_refs)
            non_caption_refs_sequential += len(seq_refs)

        print(f"INFO: Non-caption refs — chapter-based: {non_caption_refs_chapter}, sequential: {non_caption_refs_sequential}")

        if non_caption_refs_chapter >= 2 and non_caption_refs_sequential == 0:
            print(f"PASS: Component 3 — All in-text references use chapter-based format (0.2 pts)")
            total_score += 0.2
        elif non_caption_refs_chapter > 0 and non_caption_refs_sequential == 0:
            print(f"PASS: Component 3 — In-text references use chapter-based format (0.2 pts)")
            total_score += 0.2
        elif non_caption_refs_chapter > non_caption_refs_sequential:
            print(f"PARTIAL: Component 3 — Some references updated (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 3 — References still use sequential numbering")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
