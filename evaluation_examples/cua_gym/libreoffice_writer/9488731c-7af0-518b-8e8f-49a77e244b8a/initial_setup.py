"""
Initial Setup: Invoice layout with tab stops
Task ID: wrpara_037
Domain: libreoffice_writer

Creates a document with company letterhead and INVOICE heading.
The agent must then set up tab stops and add invoice line items.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Company letterhead
    company_para = doc.add_paragraph()
    company_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = company_para.add_run('Meridian Solutions Ltd.')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    address_para = doc.add_paragraph()
    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = address_para.add_run('742 Evergreen Business Park, Suite 300\nPortland, OR 97205\nTel: (503) 555-0187 | accounts@meridiansolutions.com')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Horizontal rule (simulated with bottom border)
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_after = Pt(6)
    pPr = hr_para._element.get_or_add_pPr()
    pBdr = pPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr', {})
    bottom = pBdr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom', {
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'single',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz': '6',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space': '1',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': '1A3C6E',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Invoice heading
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_para.paragraph_format.space_before = Pt(12)
    heading_para.paragraph_format.space_after = Pt(6)
    run = heading_para.add_run('INVOICE')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # Invoice metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = meta_para.add_run('Invoice No: INV-2025-0483\nDate: March 28, 2025\nDue Date: April 27, 2025')
    run.font.size = Pt(10)

    # Bill To section
    bill_para = doc.add_paragraph()
    bill_para.paragraph_format.space_before = Pt(12)
    run = bill_para.add_run('Bill To:')
    run.bold = True
    run.font.size = Pt(10)

    client_para = doc.add_paragraph()
    run = client_para.add_run('Cascade Digital Agency\n1200 NW Flanders St\nPortland, OR 97209\nAttn: Rebecca Torres')
    run.font.size = Pt(10)

    # Empty paragraph before where invoice items will go
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
