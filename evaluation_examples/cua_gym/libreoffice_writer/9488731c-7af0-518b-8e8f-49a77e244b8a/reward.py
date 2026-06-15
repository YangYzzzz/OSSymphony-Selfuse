"""
Reward Script: Invoice layout with tab stops
Task ID: wrpara_037
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Header row exists with correct tab-separated columns
  Component 2 (0.30): 10 line items with tab-separated data present
  Component 3 (0.25): Tab stops configured correctly (5 stops at specified positions)
  Component 4 (0.15): TOTAL line present with right-aligned tab at 16cm
  Component 5 (0.10): TOTAL line tab has heavy/double-underline leader
"""

import os

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

WORKDIR = '/home/user'
TASK_ID = 'wrpara_037'


def get_tab_stops(para):
    """Get non-CLEAR tab stops from a paragraph."""
    result = []
    tabs = para.paragraph_format.tab_stops
    if tabs:
        for ts in tabs:
            if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                continue
            # Skip default LEFT@0
            if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
                continue
            result.append(ts)
    return result


def cm_to_emu(cm):
    """Convert centimeters to EMU."""
    return int(cm * 360000)


def position_matches(actual_emu, expected_cm, tolerance_cm=0.3):
    """Check if a tab stop position matches expected cm within tolerance."""
    expected_emu = cm_to_emu(expected_cm)
    return abs(actual_emu - expected_emu) <= cm_to_emu(tolerance_cm)


def find_tab_paragraphs(doc, start_index=0):
    """Find paragraphs containing tab characters starting from a given index."""
    tab_paras = []
    for i in range(start_index, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if '\t' in para.text:
            tab_paras.append((i, para))
    return tab_paras


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # The initial document has 8 paragraphs (paras 0-7) as letterhead/heading.
    # The golden adds paras 8+ with invoice content.
    # Find all tab-containing paragraphs after the initial letterhead block.
    # We look for tab-containing paragraphs starting from index 7 onwards.
    tab_paras = find_tab_paragraphs(doc, start_index=7)

    # Component 1: Header row exists with tab-separated columns (0.20 points)
    # The header should contain "No.", "Description", "Price", "Qty", "Total" separated by tabs
    try:
        header_found = False
        header_para = None
        for idx, para in tab_paras:
            text = para.text.lower()
            # Check for header keywords
            has_no = 'no' in text.split('\t')[0].lower().strip('.')
            has_desc = any('desc' in col.lower() for col in text.split('\t'))
            has_price = any('price' in col.lower() for col in text.split('\t'))
            has_qty = any('qty' in col.lower() or 'quantity' in col.lower() for col in text.split('\t'))
            has_total = any('total' in col.lower() for col in text.split('\t'))

            columns = [c.strip() for c in para.text.split('\t')]
            if len(columns) >= 5 and has_desc and has_price:
                header_found = True
                header_para = (idx, para)
                print(f"PASS: Component 1 -- Header row found at para {idx}: {repr(para.text)} (0.20 pts)")
                total_score += 0.20
                break

        if not header_found:
            print(f"FAIL: Component 1 -- No header row with tab-separated columns found")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: 10 line items with tab-separated data (0.30 points)
    # Line items should have number, description, price, qty, total separated by tabs
    try:
        line_items = []
        for idx, para in tab_paras:
            text = para.text.strip()
            columns = [c.strip() for c in text.split('\t')]
            if len(columns) >= 5:
                # Check if first column is a number (line item number)
                first_col = columns[0].strip()
                if first_col.isdigit() and 1 <= int(first_col) <= 99:
                    line_items.append((idx, para))

        num_items = len(line_items)
        if num_items >= 10:
            print(f"PASS: Component 2 -- Found {num_items} line items (>=10 required) (0.30 pts)")
            total_score += 0.30
        elif num_items >= 7:
            partial = round(0.30 * (num_items / 10), 2)
            print(f"PARTIAL: Component 2 -- Found {num_items}/10 line items ({partial} pts)")
            total_score += partial
        elif num_items >= 1:
            partial = round(0.30 * (num_items / 10), 2)
            print(f"PARTIAL: Component 2 -- Found {num_items}/10 line items ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- No numbered line items found")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Tab stops configured correctly on line item paragraphs (0.25 points)
    # Expected: LEFT@1cm, LEFT@3cm, DECIMAL@12cm, CENTER@14cm, DECIMAL@16cm
    try:
        expected_tabs = [
            (WD_TAB_ALIGNMENT.LEFT, 1.0),
            (WD_TAB_ALIGNMENT.LEFT, 3.0),
            (WD_TAB_ALIGNMENT.DECIMAL, 12.0),
            (WD_TAB_ALIGNMENT.CENTER, 14.0),
            (WD_TAB_ALIGNMENT.DECIMAL, 16.0),
        ]

        # Check tab stops on line item paragraphs (or header if no line items)
        check_paras = line_items if line_items else []
        if header_para and not check_paras:
            check_paras = [header_para]

        tabs_correct_count = 0
        tabs_checked = 0

        for idx, para in check_paras[:3]:  # Check first 3 line items
            tabs = get_tab_stops(para)
            tabs_checked += 1
            if len(tabs) >= 5:
                matches = 0
                for exp_align, exp_cm in expected_tabs:
                    for ts in tabs:
                        if ts.alignment == exp_align and position_matches(ts.position, exp_cm):
                            matches += 1
                            break
                if matches >= 4:  # Allow 1 minor mismatch
                    tabs_correct_count += 1

        if tabs_checked > 0 and tabs_correct_count == tabs_checked:
            print(f"PASS: Component 3 -- Tab stops correctly configured on {tabs_correct_count}/{tabs_checked} checked paragraphs (0.25 pts)")
            total_score += 0.25
        elif tabs_correct_count > 0:
            partial = round(0.25 * (tabs_correct_count / max(tabs_checked, 1)), 2)
            print(f"PARTIAL: Component 3 -- Tab stops correct on {tabs_correct_count}/{tabs_checked} checked paragraphs ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 -- Tab stops not correctly configured. Checked {tabs_checked} paragraphs")
            if check_paras:
                sample_idx, sample_para = check_paras[0]
                sample_tabs = get_tab_stops(sample_para)
                for ts in sample_tabs:
                    print(f"  Found: alignment={ts.alignment}, position={ts.position} EMU ({ts.position/360000:.2f} cm)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: TOTAL line present with right-aligned tab at 16cm (0.15 points)
    try:
        total_line_found = False
        total_para = None
        for idx, para in tab_paras:
            text = para.text.strip()
            if text.upper().startswith('TOTAL') and '\t' in text:
                # Should NOT be a line item (no leading digit)
                columns = [c.strip() for c in text.split('\t')]
                if not columns[0].strip().isdigit():
                    # Check for right-aligned tab at ~16cm
                    tabs = get_tab_stops(para)
                    for ts in tabs:
                        if ts.alignment == WD_TAB_ALIGNMENT.RIGHT and position_matches(ts.position, 16.0):
                            total_line_found = True
                            total_para = (idx, para)
                            break
                    if not total_line_found:
                        # Even without right tab, still a total line
                        total_para = (idx, para)

        if total_line_found:
            print(f"PASS: Component 4 -- TOTAL line found with right-aligned tab at 16cm (0.15 pts)")
            total_score += 0.15
        elif total_para is not None:
            print(f"PARTIAL: Component 4 -- TOTAL line found but right-aligned tab at 16cm not confirmed (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 -- No TOTAL line with tab separator found")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: TOTAL line tab has heavy/double-underline leader (0.10 points)
    try:
        leader_correct = False
        if total_para is not None:
            idx, para = total_para
            tabs = get_tab_stops(para)
            for ts in tabs:
                if ts.alignment == WD_TAB_ALIGNMENT.RIGHT and position_matches(ts.position, 16.0):
                    # WD_TAB_LEADER.HEAVY (4) corresponds to double-underline leader
                    if ts.leader == WD_TAB_LEADER.HEAVY:
                        leader_correct = True
                        print(f"PASS: Component 5 -- TOTAL tab has HEAVY leader (double-underline) (0.10 pts)")
                        total_score += 0.10
                        break
                    elif ts.leader is not None and ts.leader != WD_TAB_LEADER.SPACES:
                        # Some leader set, partial credit
                        print(f"PARTIAL: Component 5 -- TOTAL tab has leader={ts.leader}, expected HEAVY(4) (0.05 pts)")
                        total_score += 0.05
                        leader_correct = True
                        break

        if not leader_correct:
            if total_para is not None:
                idx, para = total_para
                tabs = get_tab_stops(para)
                leaders = [(ts.alignment, ts.leader) for ts in tabs]
                print(f"FAIL: Component 5 -- TOTAL tab leader not HEAVY. Found tabs: {leaders}")
            else:
                print(f"FAIL: Component 5 -- No TOTAL line found to check leader")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
