"""
Initial Setup: Configure AutoCorrect to replace 'teh' with 'the' and 'adn' with 'and'
Task ID: osworld_writer_spell_check_autocorrect_003
Domain: libreoffice_writer

Creates an empty test document and opens it in LibreOffice Writer.
No user-level AutoCorrect customizations are set up — the agent must add them.
"""

import os
import shlex
import subprocess
import time

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_spell_check_autocorrect_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Create an empty test document
    doc = Document()

    # Add a title
    doc.add_heading('AutoCorrect Test Document', level=1)

    # Add instructions paragraph
    doc.add_paragraph(
        'This document is used to test AutoCorrect configuration. '
        'Please configure AutoCorrect to replace common typing errors.'
    )

    # Add a blank paragraph for the agent to type in
    doc.add_paragraph('')

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure any user-level autocorrect files are NOT present
    # (they would pre-configure the task, giving away the answer)
    user_autocorr_dir = f'{WORKDIR}/.config/libreoffice/4/user/autocorr'
    user_acor_file = f'{user_autocorr_dir}/acor_en-US.dat'
    if os.path.exists(user_acor_file):
        os.remove(user_acor_file)
        print(f'Removed existing user autocorrect file: {user_acor_file}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
