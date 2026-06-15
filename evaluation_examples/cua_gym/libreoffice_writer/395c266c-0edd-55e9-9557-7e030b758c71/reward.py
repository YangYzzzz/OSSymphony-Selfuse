"""
Reward verification script for writer_hr_066.
Verifies an 8-page employee onboarding packet document.
Progressive scoring 0.0 - 1.0.
"""

import os
import re
import traceback

def compute_reward():
    file_path = "/home/user/writer_hr_066.docx"

    if not os.path.exists(file_path):
        print("File not found:", file_path)
        return 0.0

    try:
        from docx import Document
    except ImportError:
        print("python-docx not available")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Failed to open document: {e}")
        return 0.0

    score = 0.0

    # Gather all text (paragraphs + table cells) for content checks
    para_text = " ".join(p.text for p in doc.paragraphs).lower()
    table_text_all = " ".join(
        cell.text.lower()
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    all_text = para_text + " " + table_text_all
    num_sections = len(doc.sections)
    num_tables = len(doc.tables)

    # If essentially blank (no meaningful text), return 0.0
    if len(para_text.strip()) < 50 and num_tables == 0:
        print("Document appears blank or nearly empty")
        print(f"REWARD: 0.0")
        return 0.0

    # =========================================================================
    # 1. Document structure - sections (0.15)
    # Expect 8 sections (one per page)
    # =========================================================================
    if num_sections >= 8:
        score += 0.15
        print(f"[1] Sections: {num_sections} >= 8 => +0.15")
    elif num_sections >= 6:
        partial = 0.15 * (num_sections / 8.0)
        score += partial
        print(f"[1] Sections: {num_sections} (partial) => +{partial:.3f}")
    else:
        print(f"[1] Sections: {num_sections} < 6 => +0.00")

    # =========================================================================
    # 2. Section headers (0.15)
    # Each section should have a distinct, non-empty header
    # =========================================================================
    headers_with_text = 0
    distinct_headers = set()
    for i, section in enumerate(doc.sections):
        try:
            hdr_text = ""
            if section.header and section.header.paragraphs:
                hdr_text = section.header.paragraphs[0].text.strip()
            if hdr_text:
                headers_with_text += 1
                distinct_headers.add(hdr_text.lower())
        except Exception:
            pass

    expected_headers = min(num_sections, 8)
    if expected_headers > 0:
        # Score based on how many sections have distinct non-empty headers
        distinct_count = len(distinct_headers)
        header_ratio = min(distinct_count / 8.0, 1.0)
        header_score = 0.15 * header_ratio
        score += header_score
        print(f"[2] Headers with text: {headers_with_text}, distinct: {distinct_count} => +{header_score:.3f}")
    else:
        print(f"[2] No sections to check headers => +0.00")

    # =========================================================================
    # 3. Page 1 - Welcome Letter (0.10)
    # Contains welcome-related text, greeting
    # =========================================================================
    welcome_keywords = ["welcome", "dear", "behalf", "team member", "onboarding", "glad", "thrilled", "excited"]
    welcome_matches = sum(1 for kw in welcome_keywords if kw in all_text)
    if welcome_matches >= 3:
        score += 0.10
        print(f"[3] Welcome letter: {welcome_matches} keywords matched => +0.10")
    elif welcome_matches >= 1:
        partial = 0.10 * (welcome_matches / 3.0)
        score += partial
        print(f"[3] Welcome letter: {welcome_matches} keywords (partial) => +{partial:.3f}")
    else:
        print(f"[3] Welcome letter: no keywords found => +0.00")

    # =========================================================================
    # 4. Page 2 - Company Overview (0.10)
    # Contains mission/values related text
    # =========================================================================
    overview_keywords = ["mission", "vision", "values", "core values", "innovation", "integrity",
                         "collaboration", "company overview", "headquarters", "founded"]
    overview_matches = sum(1 for kw in overview_keywords if kw in all_text)
    if overview_matches >= 4:
        score += 0.10
        print(f"[4] Company overview: {overview_matches} keywords matched => +0.10")
    elif overview_matches >= 1:
        partial = 0.10 * min(overview_matches / 4.0, 1.0)
        score += partial
        print(f"[4] Company overview: {overview_matches} keywords (partial) => +{partial:.3f}")
    else:
        print(f"[4] Company overview: no keywords found => +0.00")

    # =========================================================================
    # 5. Page 3 - Employee Information Form (0.10)
    # Contains a table with form fields (Name, Address, etc.)
    # =========================================================================
    form_fields = ["name", "address", "phone", "email", "date of birth", "emergency",
                   "department", "position", "social security", "start date"]
    form_score = 0.0

    # Check if there's a table with form-like fields
    form_table_found = False
    for table in doc.tables:
        table_text = " ".join(cell.text.lower() for row in table.rows for cell in row.cells)
        field_matches = sum(1 for f in form_fields if f in table_text)
        if field_matches >= 4:
            form_table_found = True
            form_score = 0.10 * min(field_matches / 6.0, 1.0)
            break

    if not form_table_found:
        # Check if form fields appear in text at all (less strict)
        text_field_matches = sum(1 for f in form_fields if f in all_text)
        if text_field_matches >= 4:
            form_score = 0.05
    score += form_score
    print(f"[5] Employee form: table={'yes' if form_table_found else 'no'} => +{form_score:.3f}")

    # =========================================================================
    # 6. Page 4 - Tax Form Instructions (0.10)
    # Contains tax/W-4/I-9 related text
    # =========================================================================
    tax_keywords = ["w-4", "i-9", "tax", "withholding", "irs", "federal income",
                    "employment eligibility", "state tax"]
    tax_matches = sum(1 for kw in tax_keywords if kw in all_text)
    if tax_matches >= 3:
        score += 0.10
        print(f"[6] Tax instructions: {tax_matches} keywords matched => +0.10")
    elif tax_matches >= 1:
        partial = 0.10 * min(tax_matches / 3.0, 1.0)
        score += partial
        print(f"[6] Tax instructions: {tax_matches} keywords (partial) => +{partial:.3f}")
    else:
        print(f"[6] Tax instructions: no keywords found => +0.00")

    # =========================================================================
    # 7. Page 5 - Benefits Enrollment (0.10)
    # Contains a table with benefits info
    # =========================================================================
    benefits_keywords = ["medical", "dental", "vision", "401(k)", "life insurance",
                         "disability", "benefits", "enrollment", "coverage", "hsa", "fsa"]
    benefits_matches = sum(1 for kw in benefits_keywords if kw in all_text)

    benefits_table_found = False
    for table in doc.tables:
        table_text = " ".join(cell.text.lower() for row in table.rows for cell in row.cells)
        ben_in_table = sum(1 for kw in ["medical", "dental", "vision", "401", "insurance",
                                         "disability", "benefit", "coverage"] if kw in table_text)
        if ben_in_table >= 3:
            benefits_table_found = True
            break

    if benefits_table_found and benefits_matches >= 4:
        score += 0.10
        print(f"[7] Benefits: table=yes, {benefits_matches} keywords => +0.10")
    elif benefits_matches >= 3:
        partial = 0.10 * min(benefits_matches / 5.0, 1.0)
        score += partial
        print(f"[7] Benefits: keywords={benefits_matches} (partial) => +{partial:.3f}")
    elif benefits_matches >= 1:
        partial = 0.05 * min(benefits_matches / 3.0, 1.0)
        score += partial
        print(f"[7] Benefits: keywords={benefits_matches} (minimal) => +{partial:.3f}")
    else:
        print(f"[7] Benefits: no content found => +0.00")

    # =========================================================================
    # 8. Page 6 - IT Setup Guide (0.05)
    # Contains IT/equipment/software related text
    # =========================================================================
    it_keywords = ["laptop", "software", "vpn", "email", "password", "it setup",
                   "equipment", "network", "it department", "help desk", "workstation"]
    it_matches = sum(1 for kw in it_keywords if kw in all_text)
    if it_matches >= 3:
        score += 0.05
        print(f"[8] IT setup: {it_matches} keywords matched => +0.05")
    elif it_matches >= 1:
        partial = 0.05 * min(it_matches / 3.0, 1.0)
        score += partial
        print(f"[8] IT setup: {it_matches} keywords (partial) => +{partial:.3f}")
    else:
        print(f"[8] IT setup: no keywords found => +0.00")

    # =========================================================================
    # 9. Page 7 - First Week Schedule (0.10)
    # Contains a table with schedule/days
    # =========================================================================
    schedule_keywords = ["monday", "tuesday", "wednesday", "thursday", "friday",
                         "schedule", "first week", "orientation", "training"]
    schedule_matches = sum(1 for kw in schedule_keywords if kw in all_text)

    schedule_table_found = False
    for table in doc.tables:
        table_text = " ".join(cell.text.lower() for row in table.rows for cell in row.cells)
        day_matches = sum(1 for d in ["monday", "tuesday", "wednesday", "thursday", "friday"]
                         if d in table_text)
        if day_matches >= 3:
            schedule_table_found = True
            break

    if schedule_table_found and schedule_matches >= 4:
        score += 0.10
        print(f"[9] Schedule: table=yes, {schedule_matches} keywords => +0.10")
    elif schedule_matches >= 3:
        partial = 0.10 * min(schedule_matches / 5.0, 1.0)
        score += partial
        print(f"[9] Schedule: keywords={schedule_matches} (partial) => +{partial:.3f}")
    elif schedule_matches >= 1:
        partial = 0.05 * min(schedule_matches / 3.0, 1.0)
        score += partial
        print(f"[9] Schedule: keywords={schedule_matches} (minimal) => +{partial:.3f}")
    else:
        print(f"[9] Schedule: no content found => +0.00")

    # =========================================================================
    # 10. Page 8 - Acknowledgment/Signature (0.05)
    # Contains acknowledgment/signature related text
    # =========================================================================
    ack_keywords = ["acknowledg", "signature", "sign below", "signing", "received",
                    "reviewed", "understand", "agree", "print name", "date:"]
    ack_matches = sum(1 for kw in ack_keywords if kw in all_text)
    if ack_matches >= 3:
        score += 0.05
        print(f"[10] Acknowledgment: {ack_matches} keywords matched => +0.05")
    elif ack_matches >= 1:
        partial = 0.05 * min(ack_matches / 3.0, 1.0)
        score += partial
        print(f"[10] Acknowledgment: {ack_matches} keywords (partial) => +{partial:.3f}")
    else:
        print(f"[10] Acknowledgment: no keywords found => +0.00")

    # Clamp score to [0.0, 1.0]
    score = round(min(max(score, 0.0), 1.0), 2)
    return score


if __name__ == "__main__":
    try:
        result = compute_reward()
    except Exception as e:
        print(f"Unexpected error: {e}")
        traceback.print_exc()
        result = 0.0
    print(f"REWARD: {result}")
