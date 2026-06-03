"""
Initial Setup: Insert company logo into document header
Task ID: writer_obj_025
Domain: libreoffice_writer

Creates:
  - /home/user/Desktop/official_letter.docx  (2-page formal letter, empty header)
  - /home/user/Desktop/logo.png              (400x300 pixel company logo)
"""

import os
import shlex
import subprocess
import time
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'official_letter'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
LOGO_PATH = f'{WORKDIR}/logo.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_logo():
    """Create a 400x300 pixel company logo PNG."""
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new('RGB', (400, 300), color=(25, 60, 120))
    draw = ImageDraw.Draw(img)

    # Draw a simple logo: white rectangle border
    draw.rectangle([20, 20, 380, 280], outline=(255, 255, 255), width=4)

    # Company initials in large text
    draw.rectangle([80, 60, 200, 160], fill=(255, 255, 255))
    draw.text((90, 70), "TC", fill=(25, 60, 120))

    # Company name area
    draw.rectangle([60, 180, 340, 240], fill=(200, 200, 200))
    draw.text((70, 190), "TechCorp International", fill=(25, 60, 120))

    img.save(LOGO_PATH)
    print(f'Logo created: {LOGO_PATH}')


def create_initial_document():
    """Create a 2-page formal letter with empty header."""
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Header: intentionally empty ---
    header = section.header
    header.is_linked_to_previous = False
    # Leave header paragraphs empty (do not add any text or image)

    # --- Page 1: Letter heading ---
    # Sender info
    sender_para = doc.add_paragraph()
    sender_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = sender_para.add_run('TechCorp International Ltd.')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    addr_para = doc.add_paragraph()
    addr_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = addr_para.add_run('Suite 400, 1200 Innovation Drive\nSilicon Valley, CA 94025\nUnited States')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = date_para.add_run('15 March 2025')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Recipient
    recip_para = doc.add_paragraph()
    run = recip_para.add_run('Ms. Angela Hoffman\nDirector of Procurement\nGlobal Supply Partners Inc.\n900 Commerce Blvd, Suite 310\nNew York, NY 10018')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Salutation
    sal = doc.add_paragraph()
    run = sal.add_run('Dear Ms. Hoffman,')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Subject line
    subj = doc.add_paragraph()
    run = subj.add_run('Re: Partnership Agreement Renewal — FY 2025-2026')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Body paragraph 1
    body1 = doc.add_paragraph()
    body1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body1.add_run(
        'I am writing on behalf of TechCorp International Ltd. to formally express our interest in renewing '
        'the strategic partnership agreement between our organizations, which is scheduled to expire on '
        '30 June 2025. Over the past three years, our collaboration has yielded significant mutual benefits, '
        'including a 28% reduction in supply chain costs and enhanced product delivery timelines.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Body paragraph 2
    body2 = doc.add_paragraph()
    body2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body2.add_run(
        'We propose scheduling a preliminary meeting during the week of 7 April 2025 to review current '
        'terms and discuss potential enhancements to our cooperative framework. In preparation, we would '
        'be pleased to share an updated proposal document outlining our suggested revisions to the service '
        'level agreements, pricing structure, and joint marketing initiatives.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Closing (page 1)
    close1 = doc.add_paragraph()
    body2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = close1.add_run(
        'Please find enclosed our preliminary renewal proposal (Enclosure A) along with updated compliance '
        'certifications. We look forward to your favourable response and to continuing our productive '
        'working relationship. Please do not hesitate to contact me directly should you require any '
        'additional information at this stage.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Page break → start page 2
    doc.add_page_break()

    # --- Page 2: Continuation / Terms Summary ---
    heading2 = doc.add_paragraph()
    run = heading2.add_run('Summary of Proposed Renewal Terms')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(13)

    doc.add_paragraph()

    # Table of proposed terms
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, heading_text in enumerate(['Term', 'Current Agreement', 'Proposed Renewal']):
        run = hdr_cells[i].paragraphs[0].add_run(heading_text)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    rows_data = [
        ('Contract Duration', '24 months', '36 months'),
        ('Annual Value', 'USD 1,200,000', 'USD 1,450,000'),
        ('Service Level', '98.5% uptime', '99.2% uptime'),
        ('Review Cycle', 'Annual', 'Semi-annual'),
        ('Penalty Clause', '2% per incident', '1.5% per incident'),
        ('Exclusivity', 'Non-exclusive', 'Preferred supplier'),
    ]
    for row_data in rows_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            run = row_cells[i].paragraphs[0].add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    doc.add_paragraph()

    # Page 2 closing
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = closing_para.add_run(
        'The above terms are indicative and subject to final negotiation. TechCorp International is committed '
        'to reaching a mutually beneficial agreement prior to the current contract expiry date. We are confident '
        'that the proposed enhancements will further strengthen our long-standing business relationship and '
        'contribute positively to both organisations\' strategic objectives for the coming fiscal year.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Signature block
    sig_para = doc.add_paragraph()
    run = sig_para.add_run('Yours sincerely,')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    name_para = doc.add_paragraph()
    run = name_para.add_run('Jonathan Reeves')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    title_para = doc.add_paragraph()
    run = title_para.add_run('Chief Partnership Officer\nTechCorp International Ltd.\nEmail: j.reeves@techcorp.com\nTel: +1 (650) 555-0192')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)

    # Save
    doc.save(OUTPUT)
    print(f'Initial document created: {OUTPUT}')


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    # Create logo PNG
    create_logo()

    # Create the initial .docx with empty header
    create_initial_document()

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
