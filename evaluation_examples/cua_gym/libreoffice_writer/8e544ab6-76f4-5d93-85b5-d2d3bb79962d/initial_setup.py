"""
Initial Setup: Python tutorial document about dictionaries — without background/padding on code paragraphs
Task ID: writer_para_036
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_para_036'
OUTPUT = f'{WORKDIR}/Desktop/python_tutorial.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = os.path.join(WORKDIR, 'Desktop')
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # Paragraph 1: Heading 1
    doc.add_heading('Python Tutorial: Working with Dictionaries', level=1)

    # Paragraph 2: Body text
    doc.add_paragraph(
        'Dictionaries are one of the most versatile data structures in Python. '
        'They store key-value pairs and provide O(1) average time complexity for lookups.'
    )

    # Paragraph 3: Sub-heading (Heading 2)
    doc.add_heading('Creating a Dictionary', level=2)

    # Paragraph 4: Code example — NO background, NO padding in initial state
    p4 = doc.add_paragraph(
        'student = {"name": "Alice", "age": 22, "major": "Computer Science", "gpa": 3.85}'
    )
    # No special formatting on this paragraph in initial state

    # Paragraph 5: Body text
    doc.add_paragraph(
        'You can access values using their keys. '
        'If the key does not exist, Python will raise a KeyError exception.'
    )

    # Paragraph 6: Sub-heading (Heading 2)
    doc.add_heading('Iterating Over a Dictionary', level=2)

    # Paragraph 7: Code example — NO background, NO padding in initial state
    p7 = doc.add_paragraph(
        'for key, value in student.items():    print(f"{key}: {value}")'
    )
    # No special formatting on this paragraph in initial state

    # Paragraph 8: Body text
    doc.add_paragraph(
        'This will print each key-value pair on a separate line.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
