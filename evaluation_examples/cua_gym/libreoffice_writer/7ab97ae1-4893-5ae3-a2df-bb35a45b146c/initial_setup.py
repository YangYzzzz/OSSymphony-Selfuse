"""
Initial Setup: Insert a cross-reference to Table 2 showing its page number
Task ID: writer_tm_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add an explicit page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def add_caption(doc, caption_text, seq_num):
    """Add a table caption using SEQ field (like LibreOffice/Word does).

    Creates: 'Table N: Caption Text' with a bookmark _Ref_Table_N around the SEQ field.
    """
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)

    # "Table " prefix
    run_prefix = para.add_run("Table ")
    run_prefix.bold = True
    run_prefix.font.size = Pt(10)

    # Bookmark start around the SEQ field
    bookmark_name = f"_Ref_Table_{seq_num}"
    bookmark_id = str(seq_num + 100)
    bm_start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="{bookmark_id}" w:name="{bookmark_name}"/>'
    )
    para._element.append(bm_start)

    # SEQ field: begin
    run_begin = para.add_run()
    run_begin.font.size = Pt(10)
    run_begin.bold = True
    fld_begin = run_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_begin._element.append(fld_begin)

    # SEQ field: instrText
    run_instr = para.add_run()
    run_instr.font.size = Pt(10)
    run_instr.bold = True
    instr_text = run_instr._element.makeelement(qn('w:instrText'), {})
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' SEQ Table \\* ARABIC '
    run_instr._element.append(instr_text)

    # SEQ field: separate
    run_sep = para.add_run()
    run_sep.font.size = Pt(10)
    run_sep.bold = True
    fld_sep = run_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run_sep._element.append(fld_sep)

    # SEQ field: cached value
    run_val = para.add_run(str(seq_num))
    run_val.font.size = Pt(10)
    run_val.bold = True

    # SEQ field: end
    run_end = para.add_run()
    run_end.font.size = Pt(10)
    run_end.bold = True
    fld_end = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_end._element.append(fld_end)

    # Bookmark end
    bm_end = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="{bookmark_id}"/>'
    )
    para._element.append(bm_end)

    # Caption text after number
    run_caption = para.add_run(f": {caption_text}")
    run_caption.bold = True
    run_caption.font.size = Pt(10)

    return bookmark_name


def add_filler_paragraphs(doc, text_lines):
    """Add multiple paragraphs of filler content."""
    for line in text_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Liberation Serif"


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Liberation Serif'
    font.size = Pt(11)

    # === PAGE 1 ===
    # Title
    heading = doc.add_heading("Quarterly Business Performance Report", level=1)
    heading.paragraph_format.space_after = Pt(12)

    add_filler_paragraphs(doc, [
        "This report provides a comprehensive overview of the company's financial performance "
        "during the third quarter of 2025. The analysis covers revenue trends, cost structures, "
        "and key performance indicators across all business divisions.",
        "",
        "The following sections present detailed breakdowns of our operational metrics, "
        "including sales figures by region, departmental expenses, and projected growth targets "
        "for the upcoming fiscal year.",
    ])

    # Table 1: Revenue Summary
    add_caption(doc, "Revenue Summary by Region", 1)

    table1 = doc.add_table(rows=8, cols=4, style="Table Grid")
    headers1 = ["Region", "Q3 Revenue ($K)", "Q2 Revenue ($K)", "Growth (%)"]
    for j, h in enumerate(headers1):
        cell = table1.cell(0, j)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data1 = [
        ["North America", "12,450", "11,200", "11.2"],
        ["Europe", "8,730", "8,100", "7.8"],
        ["Asia-Pacific", "6,290", "5,850", "7.5"],
        ["Latin America", "3,180", "2,990", "6.4"],
        ["Middle East & Africa", "2,410", "2,150", "12.1"],
        ["Oceania", "1,890", "1,760", "7.4"],
        ["Total", "34,950", "32,050", "9.0"],
    ]
    for i, row_data in enumerate(data1, 1):
        for j, val in enumerate(row_data):
            cell = table1.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 7:  # Total row
                run.bold = True

    add_filler_paragraphs(doc, [
        "",
        "As shown in Table 1, overall revenue grew by 9.0% compared to Q2 2025. "
        "The strongest growth was observed in the Middle East & Africa region at 12.1%, "
        "followed by North America at 11.2%.",
        "",
        "The cost analysis can be found on page ",
    ])

    # NOTE: The cursor should be right after "page " - the agent needs to insert
    # a cross-reference to Table 2 here showing the page number.

    add_filler_paragraphs(doc, [
        "",
        "The following sections provide more detailed breakdowns of our operating expenses "
        "and departmental budget allocations. Management has identified several areas where "
        "cost optimization strategies could yield significant savings in Q4.",
    ])

    # === PAGE 2 ===
    add_page_break(doc)

    doc.add_heading("Operational Analysis", level=2)

    add_filler_paragraphs(doc, [
        "Our operational efficiency metrics continued to improve throughout Q3 2025. "
        "The implementation of the new enterprise resource planning system has streamlined "
        "procurement processes, reducing average order processing time by 23%.",
        "",
        "Supply chain disruptions that affected Q1 and Q2 have largely been resolved. "
        "Inventory turnover rates have returned to pre-disruption levels, with warehouse "
        "utilization now at 87% across all distribution centers.",
        "",
        "Employee productivity, measured in revenue per full-time equivalent, increased "
        "by 4.7% compared to the previous quarter. This improvement is attributed to "
        "the completion of the digital transformation initiative in the sales department.",
        "",
        "Customer satisfaction scores remained stable at 4.2 out of 5.0, with notable "
        "improvements in the technical support category. Net Promoter Score increased "
        "from 42 to 47, indicating growing customer loyalty.",
        "",
        "Strategic partnerships established in Q2 with three major distributors in the "
        "Asia-Pacific region have begun yielding results, contributing approximately $820K "
        "in additional revenue during Q3.",
        "",
        "The research and development division completed two major product milestones, "
        "with the beta release of the CloudSync platform receiving positive feedback from "
        "early adopters. Full market launch is planned for Q1 2026.",
    ])

    # === PAGE 3 ===
    add_page_break(doc)

    doc.add_heading("Market Trends and Competitive Landscape", level=2)

    add_filler_paragraphs(doc, [
        "The global market for enterprise software solutions continued its upward trajectory "
        "in Q3 2025, with an estimated market size of $298 billion, up 8.3% year-over-year.",
        "",
        "Key competitive developments during the quarter included the acquisition of "
        "DataStream Inc. by GlobalTech Corp for $4.2 billion, signaling further consolidation "
        "in the cloud infrastructure segment.",
        "",
        "Emerging technologies, particularly generative AI and edge computing, are reshaping "
        "customer expectations. Our product roadmap has been updated to incorporate AI-assisted "
        "features in 60% of our product portfolio by mid-2026.",
        "",
        "Regulatory changes in the European Union, specifically the updated Digital Markets Act, "
        "have created both challenges and opportunities. Compliance costs are estimated at $1.2M "
        "annually, but the level playing field benefits our competitive positioning.",
        "",
        "Customer churn in our SMB segment decreased from 5.8% to 4.3% following the "
        "introduction of the FlexPlan pricing model. Enterprise segment retention remained "
        "strong at 97.2%.",
        "",
        "Looking ahead to Q4, we anticipate continued momentum driven by the holiday season "
        "demand surge, new product launches, and the expansion of our channel partner network "
        "in emerging markets.",
    ])

    # === PAGE 4 ===
    add_page_break(doc)

    doc.add_heading("Cost Analysis", level=2)

    add_filler_paragraphs(doc, [
        "The following table presents a detailed breakdown of operating costs by department "
        "for Q3 2025, compared against budgeted amounts and Q2 actuals.",
    ])

    # Table 2: Cost Analysis (the target table for cross-reference)
    add_caption(doc, "Cost Analysis", 2)

    table2 = doc.add_table(rows=10, cols=5, style="Table Grid")
    headers2 = ["Department", "Q3 Budget ($K)", "Q3 Actual ($K)", "Variance ($K)", "Q2 Actual ($K)"]
    for j, h in enumerate(headers2):
        cell = table2.cell(0, j)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data2 = [
        ["Engineering", "4,200", "4,350", "-150", "4,180"],
        ["Sales & Marketing", "3,800", "3,650", "150", "3,720"],
        ["Operations", "2,500", "2,480", "20", "2,510"],
        ["Human Resources", "1,200", "1,180", "20", "1,190"],
        ["Finance", "900", "870", "30", "880"],
        ["Legal & Compliance", "650", "720", "-70", "640"],
        ["IT Infrastructure", "1,800", "1,750", "50", "1,830"],
        ["R&D", "2,400", "2,520", "-120", "2,380"],
        ["Total", "17,450", "17,520", "-70", "17,330"],
    ]
    for i, row_data in enumerate(data2, 1):
        for j, val in enumerate(row_data):
            cell = table2.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 9:  # Total row
                run.bold = True

    add_filler_paragraphs(doc, [
        "",
        "Total operating costs for Q3 came in at $17,520K, slightly above the budgeted $17,450K. "
        "The primary overruns were in Engineering (+$150K) due to unplanned infrastructure upgrades "
        "and R&D (+$120K) for accelerated prototype development.",
    ])

    # === PAGE 5 (or late page 4) ===
    add_filler_paragraphs(doc, [
        "",
        "Sales & Marketing achieved a favorable variance of $150K, primarily due to the "
        "postponement of the annual brand campaign to Q4 and lower travel expenses resulting "
        "from increased use of virtual client meetings.",
    ])

    # Table 3: Budget Projections
    add_caption(doc, "Q4 Budget Projections", 3)

    table3 = doc.add_table(rows=6, cols=3, style="Table Grid")
    headers3 = ["Category", "Q4 Projected ($K)", "Annual Target ($K)"]
    for j, h in enumerate(headers3):
        cell = table3.cell(0, j)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data3 = [
        ["Revenue", "38,200", "140,000"],
        ["Operating Costs", "18,100", "70,500"],
        ["Gross Margin", "20,100", "69,500"],
        ["Capital Expenditure", "3,500", "12,800"],
        ["Net Income", "12,400", "42,200"],
    ]
    for i, row_data in enumerate(data3, 1):
        for j, val in enumerate(row_data):
            cell = table3.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if val in ["Revenue", "Net Income"]:
                run.bold = True

    add_filler_paragraphs(doc, [
        "",
        "Based on current trends and confirmed pipeline deals, management projects Q4 revenue "
        "of $38,200K, representing a 9.3% increase over Q3. This forecast assumes no major "
        "disruptions to supply chains or significant changes in the regulatory environment.",
    ])

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
