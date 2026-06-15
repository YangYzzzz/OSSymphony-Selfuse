"""
Reward Script: Insert a cross-reference to Table 2 showing the page number
Task ID: writer_tm_057
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): PAGEREF field referencing _Ref_Table_2 exists in the document
  Component 2 (0.3): The PAGEREF field is in the correct paragraph ("cost analysis can be found on page")
  Component 3 (0.3): The cached display value is a valid page number (numeric)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_057'


def persist_app_state(domain: str):
    """Best-effort save via Ctrl+S for LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that a PAGEREF cross-reference to Table 2 has been inserted,
    showing the page number where Table 2 is located.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: PAGEREF field referencing Table 2 bookmark exists (0.4 points)
    # The cross-reference should be a PAGEREF field code pointing to _Ref_Table_2
    try:
        body = doc.element.body
        instr_elements = body.findall('.//w:instrText', ns)
        pageref_table2_found = False
        for instr in instr_elements:
            if instr.text and 'PAGEREF' in instr.text and 'Table_2' in instr.text:
                pageref_table2_found = True
                print(f"PASS: Component 1 — PAGEREF field referencing Table 2 found: {instr.text.strip()!r} (0.4 pts)")
                total_score += 0.4
                break
        if not pageref_table2_found:
            # Also check for REF field with \p switch (another valid cross-reference form)
            for instr in instr_elements:
                if instr.text and 'REF' in instr.text and 'Table_2' in instr.text:
                    pageref_table2_found = True
                    print(f"PASS: Component 1 — REF field referencing Table 2 found: {instr.text.strip()!r} (0.4 pts)")
                    total_score += 0.4
                    break
        if not pageref_table2_found:
            all_instrs = [i.text.strip() for i in instr_elements if i.text]
            print(f"FAIL: Component 1 — No PAGEREF/REF field referencing Table 2. All field codes: {all_instrs}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: The PAGEREF field is in the correct paragraph (0.3 points)
    # It should be in the paragraph containing "cost analysis can be found on page"
    try:
        pageref_in_correct_para = False
        for para in doc.paragraphs:
            para_text = para.text.lower()
            if 'cost analysis' in para_text and 'page' in para_text:
                # Check if this paragraph contains a PAGEREF or REF field code
                para_instrs = para._element.findall('.//w:instrText', ns)
                for instr in para_instrs:
                    if instr.text and ('PAGEREF' in instr.text or 'REF' in instr.text) and 'Table_2' in instr.text:
                        pageref_in_correct_para = True
                        print(f"PASS: Component 2 — PAGEREF field is in the correct paragraph: '{para.text[:60]}...' (0.3 pts)")
                        total_score += 0.3
                        break
                if pageref_in_correct_para:
                    break
        if not pageref_in_correct_para:
            print(f"FAIL: Component 2 — PAGEREF field not found in the 'cost analysis...page' paragraph")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: The cached display value is a valid page number (0.3 points)
    # After the fldChar separate, there should be a run with a numeric page number
    try:
        cached_value_valid = False
        for para in doc.paragraphs:
            para_xml = para._element
            fld_chars = para_xml.findall('.//w:fldChar', ns)

            # Look for PAGEREF fields in this paragraph
            para_instrs = para_xml.findall('.//w:instrText', ns)
            has_pageref_table2 = any(
                i.text and ('PAGEREF' in i.text or 'REF' in i.text) and 'Table_2' in i.text
                for i in para_instrs
            )
            if not has_pageref_table2:
                continue

            # Find the cached value between fldChar separate and fldChar end
            runs = para_xml.findall('.//w:r', ns)
            in_field = False
            after_separate = False
            for run in runs:
                fld_char = run.find('w:fldChar', ns)
                if fld_char is not None:
                    fld_type = fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', '')
                    instr_in_run = run.find('w:instrText', ns)
                    if fld_type == 'begin':
                        in_field = True
                        after_separate = False
                    elif fld_type == 'separate':
                        after_separate = True
                    elif fld_type == 'end':
                        in_field = False
                        after_separate = False
                elif after_separate:
                    # This run contains the cached display value
                    t_elem = run.find('w:t', ns)
                    if t_elem is not None and t_elem.text is not None:
                        cached_text = t_elem.text.strip()
                        if cached_text.isdigit() and int(cached_text) > 0:
                            cached_value_valid = True
                            print(f"PASS: Component 3 — Cached page number is '{cached_text}' (0.3 pts)")
                            total_score += 0.3
                        else:
                            print(f"FAIL: Component 3 — Cached value '{cached_text}' is not a valid page number")
                        break

        if not cached_value_valid and total_score < 0.7:
            # If we didn't find via XML walking but the paragraph text ends with a digit
            # (python-docx .text includes cached field values), check that
            for para in doc.paragraphs:
                if 'cost analysis' in para.text.lower() and 'page' in para.text.lower():
                    text = para.text.strip()
                    # Check if text ends with a number after "page"
                    import re
                    match = re.search(r'page\s+(\d+)\s*$', text, re.IGNORECASE)
                    if match:
                        page_num = match.group(1)
                        # Only count if we already confirmed a PAGEREF field exists
                        if total_score >= 0.4:
                            cached_value_valid = True
                            print(f"PASS: Component 3 — Page number '{page_num}' found in paragraph text (0.3 pts)")
                            total_score += 0.3
                    break
            if not cached_value_valid:
                print(f"FAIL: Component 3 — No valid cached page number found in PAGEREF field")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
