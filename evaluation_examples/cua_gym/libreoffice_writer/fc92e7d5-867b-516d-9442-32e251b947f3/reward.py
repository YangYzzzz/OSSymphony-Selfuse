"""
Reward Script: Set right-alignment for all cells in the 'Price' column (column 3)
Task ID: writer_tbl_014
Domain: libreoffice_writer
Scoring:
  Component 1: At least one cell in column 3 is right-aligned (0.3 pts)
  Component 2: All 5 cells in column 3 are right-aligned (0.4 pts)
  Component 3: All column 3 cells are right-aligned AND cell contents are unchanged (0.3 pts)
Total: 1.0
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_014'
FILE_PATH = f'{WORKDIR}/menu_prices.docx'

# Expected column 3 cell contents (unchanged from initial state)
EXPECTED_COL3_CONTENTS = ['Price', '$3.50', '$4.75', '$2.95', '$3.25']
# Expected number of rows in the table
EXPECTED_ROWS = 5
# Expected number of columns
EXPECTED_COLS = 3


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Task: Set horizontal alignment of all cells in the 'Price' column (column 3)
    to right-aligned in the menu_prices.docx table.

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: Verify document has at least 1 table
    try:
        if len(doc.tables) < 1:
            print("CRITICAL: No tables found in document — cannot verify task")
            print("REWARD: 0.0")
            return 0.0
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows != EXPECTED_ROWS or num_cols != EXPECTED_COLS:
            print(f"WARN: Table dimensions unexpected: {num_rows}x{num_cols}, expected {EXPECTED_ROWS}x{EXPECTED_COLS}")
    except Exception as e:
        print(f"CRITICAL: Cannot access table: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: At least one cell in column 3 (index 2) is right-aligned (0.3 points)
    # The task requires aligning column 3. On initial_env all cells are LEFT, so
    # even one RIGHT-aligned cell in column 3 is a task-change indicator.
    try:
        col3_right_count = 0
        col3_alignments = []
        for ri in range(num_rows):
            cell = table.rows[ri].cells[2]
            for para in cell.paragraphs:
                alignment = para.paragraph_format.alignment
                col3_alignments.append(alignment)
                if alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    col3_right_count += 1
                    break  # count once per cell
        if col3_right_count >= 1:
            print(f"PASS: Component 1 — {col3_right_count}/{num_rows} cells in column 3 are right-aligned (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No cells in column 3 are right-aligned; alignments: {col3_alignments}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: ALL 5 cells in column 3 (index 2) are right-aligned (0.4 points)
    # This fails on initial_env (all cells are LEFT) and passes only when all 5
    # cells in column 3 have been set to RIGHT alignment.
    try:
        failed_rows = []
        for ri in range(num_rows):
            cell = table.rows[ri].cells[2]
            cell_right = any(
                para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT
                for para in cell.paragraphs
            )
            if not cell_right:
                failed_rows.append(ri)

        all_col3_right = (len(failed_rows) == 0)

        if all_col3_right:
            print(f"PASS: Component 2 — All {num_rows} cells in column 3 (Price column) are right-aligned (0.4 pts)")
            total_score += 0.4
        else:
            cell_texts = [table.rows[ri].cells[2].text.strip() for ri in failed_rows]
            print(f"FAIL: Component 2 — Rows {failed_rows} (texts: {cell_texts}) in column 3 are NOT right-aligned")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All column 3 cells are right-aligned AND cell contents are unchanged (0.3 points)
    # This compound check fails on initial_env because column 3 is NOT right-aligned there.
    # It only passes when all column 3 cells are right-aligned AND their text content
    # matches the expected values (i.e., the task did not accidentally corrupt content).
    try:
        col3_texts = [table.rows[ri].cells[2].text.strip() for ri in range(num_rows)]
        contents_match = (col3_texts == EXPECTED_COL3_CONTENTS)

        # Re-check right alignment for all column 3 cells (compound condition)
        all_right_for_comp3 = all(
            any(
                para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT
                for para in table.rows[ri].cells[2].paragraphs
            )
            for ri in range(num_rows)
        )

        if all_right_for_comp3 and contents_match:
            print(f"PASS: Component 3 — All column 3 cells are right-aligned AND contents match {EXPECTED_COL3_CONTENTS} (0.3 pts)")
            total_score += 0.3
        elif not all_right_for_comp3:
            print(f"FAIL: Component 3 — Column 3 cells not all right-aligned; cannot confirm content integrity")
        else:
            print(f"FAIL: Component 3 — Content mismatch in column 3: found {col3_texts}, expected {EXPECTED_COL3_CONTENTS}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
