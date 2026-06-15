"""
Initial Setup: Set horizontal alignment of all cells in the 'Price' column to right-aligned
Task ID: writer_tbl_014
Domain: libreoffice_writer

Creates: /home/user/Desktop/menu_prices.docx
  - A table with 5 rows x 3 columns (Item, Category, Price)
  - All cells are currently left-aligned (task asks to right-align column 3)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'menu_prices'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Table data
    headers = ['Item', 'Category', 'Price']
    rows_data = [
        ['Espresso',  'Coffee',  '$3.50'],
        ['Latte',     'Coffee',  '$4.75'],
        ['Croissant', 'Pastry',  '$2.95'],
        ['Muffin',    'Pastry',  '$3.25'],
    ]

    # Create table: 5 rows (1 header + 4 data), 3 columns
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # Fill header row — left-aligned (default)
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        para = cell.paragraphs[0]
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(header_text)
        run.bold = True

    # Fill data rows — all cells left-aligned
    for row_idx, row_data in enumerate(rows_data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.add_run(cell_text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
