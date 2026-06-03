"""
Reward Script: Find & Replace based formatting changes in Style_Guide.docx
Task ID: writer_pd_033
Domain: libreoffice_writer

Scoring rubric (5 components, summing to 1.0):
  1. Heading 3 style font name changed to Calibri          (0.20)
  2. Heading 3 style font size changed to 13pt             (0.20)
  3. Heading 3 style bold and italic enabled                (0.20)
  4. Emphasis style color changed to dark blue (#003366)    (0.25)
  5. Heading 3 count preserved (18) and Emphasis count preserved (25) (0.15)

All checks target style-level definitions, which is where the changes occur.
Run-level properties inherit from the style (None), so we verify the style objects.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_033'


def color_distance_rgb(c1, c2):
    """Euclidean distance between two RGBColor objects."""
    return sqrt((c1[0] - c2[0])**2 + (c1[1] - c2[1])**2 + (c1[2] - c2[2])**2)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Heading 3 style checks ---
    try:
        h3_style = doc.styles['Heading 3']
    except KeyError:
        print("CRITICAL: 'Heading 3' style not found in document")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Heading 3 font name is Calibri (0.20 points)
    try:
        h3_font_name = h3_style.font.name
        if h3_font_name and h3_font_name.lower() == 'calibri':
            print(f"PASS: Component 1 — Heading 3 font name is '{h3_font_name}' (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected Heading 3 font 'Calibri', found '{h3_font_name}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Heading 3 font size is 13pt (0.20 points)
    try:
        h3_font_size = h3_style.font.size
        if h3_font_size and abs(h3_font_size.pt - 13.0) < 0.5:
            print(f"PASS: Component 2 — Heading 3 font size is {h3_font_size.pt}pt (0.20 pts)")
            total_score += 0.20
        else:
            size_val = h3_font_size.pt if h3_font_size else None
            print(f"FAIL: Component 2 — Expected Heading 3 font size 13pt, found {size_val}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Heading 3 bold AND italic are True (0.20 points)
    try:
        h3_bold = h3_style.font.bold
        h3_italic = h3_style.font.italic
        if h3_bold is True and h3_italic is True:
            print(f"PASS: Component 3 — Heading 3 bold={h3_bold}, italic={h3_italic} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — Expected Heading 3 bold=True italic=True, found bold={h3_bold} italic={h3_italic}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # --- Emphasis style checks ---
    try:
        emp_style = doc.styles['Emphasis']
    except KeyError:
        print("CRITICAL: 'Emphasis' style not found in document")
        # Still report partial score from Heading 3 checks
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 4: Emphasis style color is dark blue #003366 (0.25 points)
    try:
        emp_color = emp_style.font.color.rgb
        target_color = RGBColor(0x00, 0x33, 0x66)
        if emp_color is not None:
            dist = color_distance_rgb(emp_color, target_color)
            if dist < 15.0:
                print(f"PASS: Component 4 — Emphasis color is #{emp_color} (distance={dist:.1f}) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 4 — Emphasis color #{emp_color} too far from #003366 (distance={dist:.1f})")
        else:
            print(f"FAIL: Component 4 — Emphasis color is None (inherited/not set)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Structural integrity — Heading 3 count (18) and Emphasis count (~25) preserved (0.15 points)
    # This is a compound check: both must still exist AND the style changes above must also pass.
    # We gate this on at least some task changes being present (total_score > 0) to avoid scoring preconditions.
    try:
        h3_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 3')
        emp_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run.style and run.style.name == 'Emphasis':
                    emp_count += 1

        h3_ok = (h3_count == 18)
        emp_ok = (20 <= emp_count <= 30)  # allow some tolerance

        # Only award points if task changes are also present (at least 2 of the 4 change components passed)
        if total_score >= 0.40 and h3_ok and emp_ok:
            print(f"PASS: Component 5 — H3 count={h3_count} (expected 18), Emphasis count={emp_count} (expected ~25) (0.15 pts)")
            total_score += 0.15
        elif not h3_ok or not emp_ok:
            print(f"FAIL: Component 5 — H3 count={h3_count} (expected 18), Emphasis count={emp_count} (expected ~25)")
        else:
            print(f"FAIL: Component 5 — Structural integrity check gated on task changes (score={total_score:.2f} < 0.40)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
