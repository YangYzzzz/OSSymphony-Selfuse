"""
Initial Setup: Create Style_Guide.docx with Heading 3 and Emphasis styles
Task ID: writer_pd_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def configure_heading3_style(doc):
    """Configure Heading 3 style to 12pt Arial, not bold, not italic."""
    style = doc.styles['Heading 3']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.font.bold = False
    style.font.italic = False
    style.font.color.rgb = RGBColor(0x1F, 0x49, 0x72)  # Dark teal for headings
    # Set the east-asia font too
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = style.element.makeelement(qn('w:rPr'), {})
        style.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'Arial')
    rfonts.set(qn('w:hAnsi'), 'Arial')
    rfonts.set(qn('w:cs'), 'Arial')


def configure_emphasis_style(doc):
    """Configure Emphasis character style with default (black) color."""
    style = doc.styles['Emphasis']
    style.font.italic = True
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black (default)


def add_body_text(doc, text):
    """Add a body paragraph."""
    para = doc.add_paragraph(text)
    return para


def add_emphasis_run(para, text):
    """Add a run with Emphasis character style to an existing paragraph."""
    run = para.add_run(text)
    run.style = para.part.document.styles['Emphasis']
    return run


def add_heading3(doc, text):
    """Add a Heading 3 paragraph."""
    return doc.add_heading(text, level=3)


def create_initial():
    doc = Document()

    # Configure styles
    configure_heading3_style(doc)
    configure_emphasis_style(doc)

    # ===== Title Page =====
    title = doc.add_heading('Corporate Style Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x72)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version_para.add_run('Version 4.2 — March 2026')
    run.font.size = Pt(11)

    doc.add_page_break()

    # ===== Table of Contents Page =====
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Brand Identity and Visual Standards',
        '2. Typography Guidelines',
        '3. Color Palette and Usage',
        '4. Document Formatting Rules',
        '5. Digital Communications',
        '6. Print Materials',
        '7. Presentation Standards',
        '8. Email and Correspondence',
        '9. Social Media Guidelines',
        '10. Accessibility Requirements',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # Track emphasis count
    emphasis_count = 0

    # ===== Section 1: Brand Identity (pages 3-4) =====
    doc.add_heading('1. Brand Identity and Visual Standards', level=1)

    # H3 #1
    add_heading3(doc, '1.1 Logo Usage Requirements')
    p = add_body_text(doc, 'The Meridian Technologies logo must be displayed prominently on all ')
    add_emphasis_run(p, 'official corporate communications')  # emphasis 1
    p.add_run('. The minimum clear space around the logo should be 0.5 inches on all sides.')
    emphasis_count += 1

    add_body_text(doc, 'When placing the logo on colored backgrounds, always use the reverse-white version for dark backgrounds with a minimum contrast ratio of 4.5:1. The primary logo should be used on white or light gray backgrounds only.')

    p = add_body_text(doc, 'Logo files are available in SVG, PNG, and EPS formats from the ')
    add_emphasis_run(p, 'Brand Resources Portal')  # emphasis 2
    p.add_run(' at brand.meridiantech.com. Never recreate the logo manually.')
    emphasis_count += 1

    # H3 #2
    add_heading3(doc, '1.2 Brand Voice and Tone')
    p = add_body_text(doc, 'All written communications should maintain a ')
    add_emphasis_run(p, 'professional yet approachable')  # emphasis 3
    p.add_run(' tone. Avoid overly technical jargon when addressing non-technical audiences. Use active voice and concise sentences.')
    emphasis_count += 1

    add_body_text(doc, 'The brand voice should convey expertise, reliability, and innovation. When writing for marketing materials, focus on customer benefits rather than feature lists. For technical documentation, clarity and precision take priority over stylistic considerations.')

    doc.add_page_break()

    # ===== Section 2: Typography (pages 5-6) =====
    doc.add_heading('2. Typography Guidelines', level=1)

    # H3 #3
    add_heading3(doc, '2.1 Primary Typeface Selection')
    p = add_body_text(doc, 'The primary corporate typeface is ')
    add_emphasis_run(p, 'Calibri')  # emphasis 4
    p.add_run(' for digital documents and Arial for legacy systems. All body text should use 11pt for readability in standard documents.')
    emphasis_count += 1

    add_body_text(doc, 'For headings, the hierarchy follows: Title at 26pt, Heading 1 at 20pt, Heading 2 at 16pt, and Heading 3 at 12pt. Consistent sizing ensures visual hierarchy across all document types produced by the organization.')

    # H3 #4
    add_heading3(doc, '2.2 Font Pairing Rules')
    p = add_body_text(doc, 'When combining fonts in a single document, limit selections to ')
    add_emphasis_run(p, 'no more than three typefaces')  # emphasis 5
    p.add_run('. The recommended pairings include Calibri with Cambria for formal documents and Arial with Georgia for web-based content.')
    emphasis_count += 1

    add_body_text(doc, 'Decorative or display fonts should only be used in marketing materials and must receive prior approval from the brand team. Script fonts are prohibited in all corporate documents regardless of context.')

    # H3 #5
    add_heading3(doc, '2.3 International Character Support')
    add_body_text(doc, 'Documents intended for international distribution must use Unicode-compatible fonts. Verify character rendering for CJK (Chinese, Japanese, Korean) text blocks when applicable. The fallback font for unsupported glyphs is Noto Sans.')

    doc.add_page_break()

    # ===== Section 3: Color Palette (pages 7-8) =====
    doc.add_heading('3. Color Palette and Usage', level=1)

    # H3 #6
    add_heading3(doc, '3.1 Primary Brand Colors')
    p = add_body_text(doc, 'The primary palette consists of ')
    add_emphasis_run(p, 'Meridian Blue (#1F4972)')  # emphasis 6
    p.add_run(', Silver Gray (#A8B4C0), and Pure White (#FFFFFF). These colors must appear on all branded materials as the dominant palette.')
    emphasis_count += 1

    add_body_text(doc, 'Meridian Blue is used for headers, primary buttons, and key interactive elements. Silver Gray serves as a secondary accent and divider color. White provides the base background for all standard documents.')

    # H3 #7
    add_heading3(doc, '3.2 Secondary and Accent Colors')
    p = add_body_text(doc, 'Secondary colors include Teal (#2A9D8F), Coral (#E76F51), and Amber (#F4A261). These should be used ')
    add_emphasis_run(p, 'sparingly and only for data visualization')  # emphasis 7
    p.add_run(' or call-to-action elements.')
    emphasis_count += 1

    add_body_text(doc, 'Never use secondary colors for body text or headings. The total area coverage of secondary colors should not exceed 15% of any single page layout.')

    # H3 #8
    add_heading3(doc, '3.3 Color Accessibility Standards')
    p = add_body_text(doc, 'All text must meet ')
    add_emphasis_run(p, 'WCAG 2.1 Level AA contrast requirements')  # emphasis 8
    p.add_run('. This means a minimum contrast ratio of 4.5:1 for normal text and 3:1 for large text (18pt or 14pt bold).')
    emphasis_count += 1

    add_body_text(doc, 'Avoid conveying information through color alone. Always provide secondary indicators such as patterns, labels, or icons alongside color-coded elements. Test all color combinations using an accessibility checker before publication.')

    doc.add_page_break()

    # ===== Section 4: Document Formatting (pages 9-10) =====
    doc.add_heading('4. Document Formatting Rules', level=1)

    # H3 #9
    add_heading3(doc, '4.1 Margins and Page Layout')
    add_body_text(doc, 'Standard documents use 1-inch margins on all sides with A4 paper size for international distribution and Letter size for North American documents. Headers begin at 0.5 inches from the top edge.')

    p = add_body_text(doc, 'Landscape orientation is permitted only for ')
    add_emphasis_run(p, 'wide tables and data-heavy appendices')  # emphasis 9
    p.add_run('. Mixed orientation within a single document requires section breaks with clearly labeled page transitions.')
    emphasis_count += 1

    # H3 #10
    add_heading3(doc, '4.2 Paragraph Spacing and Indentation')
    add_body_text(doc, 'Body paragraphs use 6pt spacing after each paragraph with no first-line indentation. List items use 3pt spacing after each item with 0.25-inch hanging indent for wrapped lines.')

    p = add_body_text(doc, 'Block quotations should be indented ')
    add_emphasis_run(p, '0.5 inches from both margins')  # emphasis 10
    p.add_run(' with 10pt italic text and a left border in Silver Gray.')
    emphasis_count += 1

    add_body_text(doc, 'Avoid double spacing between paragraphs by pressing Enter twice. Use the paragraph spacing controls in the Paragraph dialog box or the styles panel to ensure consistent spacing throughout the document.')

    doc.add_page_break()

    # ===== Section 5: Digital Communications (pages 11-12) =====
    doc.add_heading('5. Digital Communications', level=1)

    # H3 #11
    add_heading3(doc, '5.1 Email Signature Standards')
    p = add_body_text(doc, 'All employee email signatures must follow the ')
    add_emphasis_run(p, 'approved signature template')  # emphasis 11
    p.add_run(' available on the intranet. Signatures should include: full name, title, department, phone number, and the Meridian Technologies logo.')
    emphasis_count += 1

    add_body_text(doc, 'Personal quotes, animated GIFs, or custom graphics are not permitted in email signatures. Social media icons may be included only for official corporate accounts managed by the marketing department.')

    # H3 #12
    add_heading3(doc, '5.2 Newsletter and Bulletin Formatting')
    p = add_body_text(doc, 'Internal newsletters use a ')
    add_emphasis_run(p, 'two-column layout with 0.3-inch gutter')  # emphasis 12
    p.add_run('. The masthead displays the Meridian Blue background with white text. Article headlines use Heading 2 style.')
    emphasis_count += 1

    add_body_text(doc, 'External bulletins require approval from the communications team before distribution. Maximum length for a single bulletin is 4 pages. Include a table of contents for bulletins exceeding 2 pages.')

    doc.add_page_break()

    # ===== Section 6: Print Materials (pages 13-14) =====
    doc.add_heading('6. Print Materials', level=1)

    # H3 #13
    add_heading3(doc, '6.1 Business Card Specifications')
    p = add_body_text(doc, 'Business cards measure 3.5 x 2 inches with ')
    add_emphasis_run(p, 'rounded corners at 0.125-inch radius')  # emphasis 13
    p.add_run('. Front side displays name, title, and contact details. Back side features the corporate logo centered on Meridian Blue background.')
    emphasis_count += 1

    add_body_text(doc, 'Card stock must be 350gsm coated matte finish. Text should be set at 8pt minimum for readability. QR codes linking to digital business cards are optional and should be placed in the bottom-right corner of the front side.')

    # H3 #14
    add_heading3(doc, '6.2 Brochure and Flyer Guidelines')
    p = add_body_text(doc, 'Tri-fold brochures use the ')
    add_emphasis_run(p, 'standard panel sizing of 3.69 x 8.5 inches')  # emphasis 14
    p.add_run('. Cover panel features a hero image with overlay text in white. Interior panels use alternating white and light gray backgrounds.')
    emphasis_count += 1

    add_body_text(doc, 'All print materials must include bleed area of 0.125 inches beyond the trim edge. CMYK color mode is mandatory for all print-ready files. RGB files will be rejected by the print vendor.')

    doc.add_page_break()

    # ===== Section 7: Presentations (pages 15-16) =====
    doc.add_heading('7. Presentation Standards', level=1)

    # H3 #15
    add_heading3(doc, '7.1 Slide Master Templates')
    p = add_body_text(doc, 'All presentations must use the ')
    add_emphasis_run(p, 'official Meridian slide deck template')  # emphasis 15
    p.add_run(' available on SharePoint. Templates include title slide, content slide, section divider, comparison layout, and closing slide formats.')
    emphasis_count += 1

    add_body_text(doc, 'Custom slide layouts may be created for specialized presentations but must maintain the corporate color palette and typography standards. Submit custom layouts for brand review before external use.')

    # H3 #16
    add_heading3(doc, '7.2 Data Visualization in Presentations')
    p = add_body_text(doc, 'Charts and graphs should use the ')
    add_emphasis_run(p, 'sequential color scheme')  # emphasis 16
    p.add_run(' starting with Meridian Blue and progressing through lighter tints. Avoid using more than 6 colors in a single chart.')
    emphasis_count += 1

    add_body_text(doc, 'Pie charts are discouraged for more than 5 categories. Use horizontal bar charts for comparison data and line charts for trend analysis. All charts must include a title, axis labels, and data source citation.')

    doc.add_page_break()

    # ===== Section 8: Email and Correspondence (pages 17-18) =====
    doc.add_heading('8. Email and Correspondence', level=1)

    # H3 #17
    add_heading3(doc, '8.1 Formal Letter Layout')
    p = add_body_text(doc, 'Formal correspondence uses ')
    add_emphasis_run(p, 'modified block format')  # emphasis 17
    p.add_run(' with the date and closing aligned to the center. Company letterhead is required for all external correspondence.')
    emphasis_count += 1

    add_body_text(doc, 'Letters exceeding one page must include a header on continuation pages showing the recipient name, date, and page number. Use "Enc." notation for enclosed documents and "cc:" for carbon copy recipients.')

    p = add_body_text(doc, 'The salutation should use the recipient\'s ')
    add_emphasis_run(p, 'preferred title and surname')  # emphasis 18
    p.add_run(' unless a first-name basis has been established. Close formal letters with "Sincerely" or "Best regards" followed by a handwritten signature space of at least 1.5 inches.')
    emphasis_count += 1

    doc.add_page_break()

    # ===== Section 9: Social Media (pages 19) =====
    doc.add_heading('9. Social Media Guidelines', level=1)

    # H3 #18
    add_heading3(doc, '9.1 Platform-Specific Formatting')
    p = add_body_text(doc, 'Each social media platform has ')
    add_emphasis_run(p, 'unique image dimension requirements')  # emphasis 19
    p.add_run('. LinkedIn posts use 1200x627 pixels, Twitter/X posts use 1600x900 pixels, and Instagram posts use 1080x1080 pixels.')
    emphasis_count += 1

    add_body_text(doc, 'Hashtags should be limited to 3-5 per post on LinkedIn and Twitter. Instagram allows up to 15 hashtags placed in the first comment rather than the caption. All social media graphics must include the Meridian watermark at 30% opacity.')

    p = add_body_text(doc, 'Video content should include ')
    add_emphasis_run(p, 'closed captions and alt text')  # emphasis 20
    p.add_run(' for accessibility. Maximum video length is 60 seconds for Instagram Reels, 2 minutes for Twitter, and 10 minutes for LinkedIn.')
    emphasis_count += 1

    add_body_text(doc, 'User-generated content may be reshared only with explicit written permission from the original creator. Tag the creator in all reshared posts and add a "Shared with permission" note.')

    doc.add_page_break()

    # ===== Section 10: Accessibility (page 20) =====
    doc.add_heading('10. Accessibility Requirements', level=1)

    p = add_body_text(doc, 'All documents must comply with ')
    add_emphasis_run(p, 'Section 508 and WCAG 2.1 guidelines')  # emphasis 21
    p.add_run('. This includes proper heading hierarchy, alt text for images, and sufficient color contrast.')
    emphasis_count += 1

    add_body_text(doc, 'Use the built-in Accessibility Checker in Microsoft Office or LibreOffice before finalizing any document. Address all errors and review all warnings before distribution.')

    p = add_body_text(doc, 'Tables must include ')
    add_emphasis_run(p, 'header row designation')  # emphasis 22
    p.add_run(' for screen reader compatibility. Avoid merged cells when possible as they create navigation difficulties for assistive technology users.')
    emphasis_count += 1

    add_body_text(doc, 'PDF documents exported from office suites must be tagged PDFs with proper reading order. Use the PDF/UA standard for documents that will be publicly distributed.')

    p = add_body_text(doc, 'Multimedia content requires ')
    add_emphasis_run(p, 'audio descriptions and transcripts')  # emphasis 23
    p.add_run(' in addition to closed captions. Provide alternative formats (large print, braille-ready) upon request.')
    emphasis_count += 1

    add_body_text(doc, 'Interactive elements in digital documents must be keyboard-navigable. Form fields should include descriptive labels and clear error messages for validation failures.')

    p = add_body_text(doc, 'Annual accessibility audits are conducted by ')
    add_emphasis_run(p, 'the Compliance and Standards team')  # emphasis 24
    p.add_run(' in Q2 of each fiscal year. Departments must remediate flagged issues within 30 business days of the audit report.')
    emphasis_count += 1

    p = add_body_text(doc, 'Training on accessible document creation is ')
    add_emphasis_run(p, 'mandatory for all content creators')  # emphasis 25
    p.add_run(' and must be renewed annually through the corporate learning management system.')
    emphasis_count += 1

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Total Emphasis spans: {emphasis_count}')

    # Count heading 3 paragraphs
    h3_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 3')
    print(f'Total Heading 3 paragraphs: {h3_count}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
