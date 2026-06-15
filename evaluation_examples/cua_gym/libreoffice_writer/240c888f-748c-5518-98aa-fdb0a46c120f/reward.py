"""
Reward Script: Insert formatted appendix section with separator page, custom headers/footers, and A-N page numbering
Task ID: writer_rd_093
Domain: libreoffice_writer
Scoring:
  Component 1: Document has multiple sections (>1) indicating appendix was added (0.15)
  Component 2: Separator page with 'APPENDIX' centered in large bold text (0.25)
  Component 3: Appendix section header shows 'Appendix' (not main doc title) (0.15)
  Component 4: Appendix footer has 'A-' prefix with page number field (0.25)
  Component 5: Appendix page numbering restarted at 1 (0.10)
  Component 6: Appendix contains placeholder content (supplementary tables/figures) (0.10)
"""

import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_093'


def persist_app_state(domain):
    """Attempt to save any unsaved LibreOffice changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify appendix insertion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = doc.sections
    num_sections = len(sections)

    # Component 1: Document has multiple sections indicating appendix was added (0.15 points)
    # Initial doc has only 1 section; golden should have 3 (main, separator, appendix)
    try:
        if num_sections >= 3:
            print(f"PASS: Component 1 — Document has {num_sections} sections (>= 3) (0.15 pts)")
            total_score += 0.15
        elif num_sections == 2:
            print(f"PARTIAL: Component 1 — Document has 2 sections, expected >= 3 (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 1 — Document has only {num_sections} section(s), expected >= 3")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Separator page with 'APPENDIX' centered in large bold text (0.25 points)
    try:
        found_appendix_separator = False
        appendix_centered = False
        appendix_bold = False
        appendix_large = False

        for para in doc.paragraphs:
            text = para.text.strip().upper()
            if text == 'APPENDIX':
                found_appendix_separator = True
                # Check centering
                alignment = para.paragraph_format.alignment
                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    appendix_centered = True
                # Check bold and size from runs
                for run in para.runs:
                    if run.font.bold:
                        appendix_bold = True
                    if run.font.size and run.font.size.pt >= 30:
                        appendix_large = True
                break

        if found_appendix_separator and appendix_centered and appendix_bold and appendix_large:
            print(f"PASS: Component 2 — 'APPENDIX' separator found: centered, bold, large text (0.25 pts)")
            total_score += 0.25
        elif found_appendix_separator and appendix_centered:
            print(f"PARTIAL: Component 2 — 'APPENDIX' found and centered but missing bold={appendix_bold} or large={appendix_large} (0.15 pts)")
            total_score += 0.15
        elif found_appendix_separator:
            print(f"PARTIAL: Component 2 — 'APPENDIX' text found but not centered (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 2 — No 'APPENDIX' separator paragraph found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Appendix section header shows 'Appendix' not main document title (0.15 points)
    # The last section (or a section after the separator) should have header text containing 'Appendix'
    # and NOT containing the main document title
    try:
        appendix_header_found = False
        main_title = ""
        # Get main title from section 0 header
        if num_sections >= 1:
            sec0_header = sections[0].header
            if sec0_header.paragraphs:
                main_title = sec0_header.paragraphs[0].text.strip()

        # Check sections after the first for an appendix-style header
        for sec_idx in range(1, num_sections):
            sec = sections[sec_idx]
            header = sec.header
            if header.is_linked_to_previous:
                continue
            if header.paragraphs:
                h_text = header.paragraphs[0].text.strip().lower()
                if 'appendix' in h_text and (not main_title or main_title.lower() not in h_text):
                    appendix_header_found = True
                    print(f"PASS: Component 3 — Section {sec_idx} header contains 'Appendix': '{header.paragraphs[0].text.strip()}' (0.15 pts)")
                    total_score += 0.15
                    break

        if not appendix_header_found:
            print(f"FAIL: Component 3 — No appendix section header with 'Appendix' text found (distinct from main title)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Appendix footer has 'A-' prefix with PAGE field code (0.25 points)
    try:
        a_prefix_found = False
        page_field_found = False
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        for sec_idx in range(1, num_sections):
            sec = sections[sec_idx]
            footer = sec.footer
            if footer.is_linked_to_previous:
                continue
            for p in footer.paragraphs:
                # Check for A- prefix in text runs
                for run in p.runs:
                    if 'A-' in run.text:
                        a_prefix_found = True
                # Check for PAGE field code in footer XML
                for instr in p._element.findall('.//w:instrText', ns):
                    if instr.text and 'PAGE' in instr.text:
                        page_field_found = True

            if a_prefix_found:
                break

        if a_prefix_found and page_field_found:
            print(f"PASS: Component 4 — Footer has 'A-' prefix with PAGE field (0.25 pts)")
            total_score += 0.25
        elif a_prefix_found:
            print(f"PARTIAL: Component 4 — Footer has 'A-' prefix but no PAGE field found (0.12 pts)")
            total_score += 0.12
        elif page_field_found:
            print(f"PARTIAL: Component 4 — Footer has PAGE field but missing 'A-' prefix (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 4 — No appendix footer with 'A-' prefix and PAGE field found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Appendix page numbering restarted at 1 (0.10 points)
    # Check w:pgNumType with w:start="1" in the appendix section
    try:
        pg_restart_found = False
        for sec_idx in range(1, num_sections):
            sec = sections[sec_idx]
            sectPr = sec._sectPr
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                start_val = pgNumType.get(qn('w:start'))
                if start_val == '1':
                    pg_restart_found = True
                    print(f"PASS: Component 5 — Section {sec_idx} has page numbering restart at 1 (0.10 pts)")
                    total_score += 0.10
                    break

        if not pg_restart_found:
            print(f"FAIL: Component 5 — No section found with page numbering restarted at 1")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Appendix contains placeholder content for supplementary tables/figures (0.10 points)
    # Check for paragraphs mentioning appendix content like 'Table', 'Figure', 'Supplementary'
    try:
        appendix_content_keywords = 0
        # Find paragraphs after the APPENDIX separator
        in_appendix = False
        for para in doc.paragraphs:
            if para.text.strip().upper() == 'APPENDIX':
                in_appendix = True
                continue
            if in_appendix:
                text_lower = para.text.lower()
                if any(kw in text_lower for kw in ['table', 'figure', 'supplementary', 'appendix a', 'appendix b', 'appendix c']):
                    appendix_content_keywords += 1

        if appendix_content_keywords >= 3:
            print(f"PASS: Component 6 — Appendix has {appendix_content_keywords} content references (tables/figures/supplementary) (0.10 pts)")
            total_score += 0.10
        elif appendix_content_keywords >= 1:
            print(f"PARTIAL: Component 6 — Appendix has {appendix_content_keywords} content references, expected >= 3 (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No appendix content found after separator page")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
