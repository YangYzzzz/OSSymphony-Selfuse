"""
Initial Setup: Create a 20-page research report document
Task ID: writer_rd_093
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_093'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section, prefix_text="Page "):
    """Add page number field to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run0 = fp.add_run(prefix_text)
    run0.font.size = Pt(10)
    # PAGE field
    r1 = fp.add_run()
    r1.font.size = Pt(10)
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)
    r2 = fp.add_run()
    r2.font.size = Pt(10)
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3.font.size = Pt(10)
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Header: Document title ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hr = hp.add_run("Advances in Neural Network Architectures for Climate Modeling")
    hr.font.size = Pt(10)
    hr.italic = True

    # --- Footer: Page numbers (Arabic) ---
    add_page_number_footer(section, "Page ")

    # --- Title Page ---
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_heading("Advances in Neural Network Architectures for Climate Modeling", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = subtitle.add_run("A Comprehensive Research Report")
    sr.font.size = Pt(16)
    sr.italic = True

    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ar = authors.add_run(
        "Dr. Elena Vasquez, Dr. James Thornton, Dr. Mei-Ling Wu\n"
        "Institute for Computational Earth Sciences\n"
        "March 2025"
    )
    ar.font.size = Pt(12)

    doc.add_page_break()

    # --- Table of Contents placeholder (page 2) ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("1. Introduction", "3"),
        ("2. Literature Review", "5"),
        ("3. Methodology", "8"),
        ("4. Neural Network Architectures", "10"),
        ("5. Experimental Results", "13"),
        ("6. Climate Projection Analysis", "15"),
        ("7. Discussion", "17"),
        ("8. Conclusions", "19"),
        ("References", "20"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
        r = p.add_run(f"{item}\t{page}")
        r.font.size = Pt(11)

    doc.add_page_break()

    # --- Chapter 1: Introduction (pages 3-4) ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Climate modeling represents one of the most computationally demanding scientific endeavors "
        "of the twenty-first century. As global temperatures continue to rise at unprecedented rates, "
        "the need for accurate, high-resolution climate projections has never been more urgent. "
        "Traditional general circulation models (GCMs), while foundational to our understanding of "
        "atmospheric dynamics, face inherent limitations in capturing the complex, nonlinear "
        "interactions that drive climate variability across multiple spatial and temporal scales."
    )
    doc.add_paragraph(
        "Recent advances in deep learning have opened new avenues for enhancing climate prediction "
        "accuracy. Neural network architectures, particularly those designed for sequential and "
        "spatial data processing, have demonstrated remarkable capabilities in pattern recognition "
        "and feature extraction from high-dimensional datasets. The intersection of these "
        "computational approaches with climate science presents transformative opportunities for "
        "improving both short-term weather forecasting and long-term climate projections."
    )
    doc.add_heading("1.1 Research Objectives", level=2)
    doc.add_paragraph(
        "This report presents a systematic investigation into the application of advanced neural "
        "network architectures for regional climate modeling. Our primary objectives include: "
        "(a) evaluating the performance of transformer-based models against traditional convolutional "
        "approaches for temperature and precipitation prediction; (b) developing a hybrid architecture "
        "that leverages both physical constraints and data-driven learning; and (c) assessing the "
        "computational efficiency trade-offs between model complexity and prediction accuracy."
    )
    doc.add_heading("1.2 Scope and Significance", level=2)
    doc.add_paragraph(
        "The scope of this research encompasses both global circulation patterns and regional "
        "downscaling applications. We focus particularly on the North Atlantic Oscillation (NAO) "
        "and its cascading effects on European climate variability. The significance of this work "
        "extends beyond pure scientific inquiry, as improved climate models directly inform policy "
        "decisions regarding infrastructure planning, agricultural adaptation, and disaster "
        "preparedness strategies across multiple nations and economic sectors."
    )

    doc.add_page_break()

    # --- Chapter 2: Literature Review (pages 5-7) ---
    doc.add_heading("2. Literature Review", level=1)
    doc.add_heading("2.1 Traditional Climate Models", level=2)
    doc.add_paragraph(
        "The development of numerical weather prediction dates back to the pioneering work of "
        "Vilhelm Bjerknes in the early 1900s and the computational implementations by Lewis Fry "
        "Richardson. Modern GCMs, including the Community Earth System Model (CESM) and the "
        "Geophysical Fluid Dynamics Laboratory (GFDL) model, solve primitive equations governing "
        "atmospheric and oceanic circulation on discrete grids typically ranging from 50 to 200 km "
        "resolution. While these models have steadily improved through the Coupled Model "
        "Intercomparison Project (CMIP) phases, systematic biases persist, particularly in "
        "representing cloud microphysics, convective processes, and land-atmosphere feedbacks "
        "(Henderson et al., 2023; Ramirez and Okonkwo, 2024)."
    )
    doc.add_heading("2.2 Machine Learning in Climate Science", level=2)
    doc.add_paragraph(
        "The application of machine learning to climate science has evolved rapidly since the "
        "initial studies by Krasnopolsky and Fox-Rabinovitz (2006) on neural network parameterizations. "
        "Convolutional neural networks (CNNs) were first applied to climate downscaling by "
        "Vandal et al. (2017), achieving superior performance to traditional statistical methods. "
        "Subsequent work by Nguyen et al. (2022) demonstrated that attention mechanisms could "
        "capture long-range spatial dependencies crucial for accurate precipitation forecasting. "
        "The FourCastNet architecture (Pathak et al., 2022) achieved forecast skill comparable "
        "to operational numerical models at a fraction of the computational cost."
    )
    doc.add_heading("2.3 Transformer Architectures for Spatiotemporal Data", level=2)
    doc.add_paragraph(
        "Transformer models, originally developed for natural language processing, have shown "
        "remarkable adaptability to spatiotemporal prediction tasks. The Vision Transformer (ViT) "
        "and its variants treat spatial patches as token sequences, enabling self-attention "
        "mechanisms to learn complex spatial relationships. Pangu-Weather (Bi et al., 2023) and "
        "GraphCast (Lam et al., 2023) represent state-of-the-art applications of transformer-like "
        "architectures to global weather forecasting, achieving unprecedented accuracy at lead "
        "times of up to 10 days."
    )
    doc.add_paragraph(
        "However, critical gaps remain in the literature regarding the application of these "
        "architectures to decadal-scale climate projections, where the stationarity assumptions "
        "underlying many machine learning approaches may not hold. Furthermore, the integration "
        "of physical conservation laws into transformer architectures remains an active area of "
        "research, with approaches ranging from soft penalty terms to hard constraint layers "
        "(Beucler et al., 2024; Kim and Park, 2024)."
    )

    doc.add_page_break()

    # --- Chapter 3: Methodology (pages 8-9) ---
    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Data Sources and Preprocessing", level=2)
    doc.add_paragraph(
        "Our study utilizes ERA5 reanalysis data from the European Centre for Medium-Range "
        "Weather Forecasts (ECMWF), spanning the period 1979-2024 at 0.25-degree spatial "
        "resolution. The dataset encompasses 37 pressure levels and includes variables such as "
        "temperature, geopotential height, specific humidity, and wind components. Supplementary "
        "observational data from the Global Historical Climatology Network (GHCN) and the "
        "Automated Surface Observing System (ASOS) were used for validation purposes."
    )
    doc.add_heading("3.2 Model Architecture Design", level=2)
    doc.add_paragraph(
        "We propose a hybrid architecture, termed ClimateFormer, that combines a spatial encoder "
        "based on Swin Transformer blocks with a temporal decoder employing causal attention "
        "mechanisms. The spatial encoder processes atmospheric fields at multiple resolutions "
        "through a hierarchical patch merging strategy, while the temporal decoder autoregressively "
        "generates future states conditioned on physical constraint embeddings derived from "
        "conservation of mass, energy, and momentum."
    )
    doc.add_heading("3.3 Training Procedure", level=2)
    doc.add_paragraph(
        "Models were trained on the NVIDIA DGX SuperPOD cluster using 64 A100 GPUs with mixed-precision "
        "training. The training schedule employed a cosine annealing learning rate with warm restarts, "
        "starting at 3e-4 and decaying to 1e-6 over 200 epochs. Data augmentation included random "
        "temporal shifts, meridional flipping, and Gaussian noise injection to improve generalization. "
        "The loss function combines a weighted mean squared error term for primary variables with a "
        "physics-informed penalty term enforcing conservation laws within each prediction step."
    )
    doc.add_heading("3.4 Evaluation Metrics", level=2)
    doc.add_paragraph(
        "Model performance was assessed using Root Mean Square Error (RMSE), Anomaly Correlation "
        "Coefficient (ACC), and Continuous Ranked Probability Score (CRPS) against ERA5 reanalysis "
        "and independent station observations. Skill scores were computed relative to climatological "
        "baselines and persistence forecasts. Additionally, we evaluated the models' ability to "
        "reproduce known teleconnection patterns and extreme event statistics using the Kolmogorov-Smirnov "
        "test and quantile-quantile analysis."
    )

    doc.add_page_break()

    # --- Chapter 4: Neural Network Architectures (pages 10-12) ---
    doc.add_heading("4. Neural Network Architectures", level=1)
    doc.add_heading("4.1 Baseline CNN Model", level=2)
    doc.add_paragraph(
        "The baseline convolutional architecture follows a U-Net design with residual connections, "
        "processing input atmospheric fields at 64x128 grid resolution. The encoder consists of "
        "four downsampling blocks, each comprising two 3x3 convolutional layers with batch "
        "normalization and GELU activation, followed by 2x2 max pooling. The decoder mirrors "
        "this structure with transposed convolutions for upsampling. Skip connections between "
        "corresponding encoder and decoder levels preserve fine-grained spatial information."
    )
    doc.add_heading("4.2 ClimateFormer Architecture", level=2)
    doc.add_paragraph(
        "The ClimateFormer architecture represents our primary contribution. The model processes "
        "atmospheric fields as sequences of overlapping patches (patch size 8x8, stride 4), "
        "projected into a 768-dimensional embedding space. The spatial encoder consists of 12 "
        "Swin Transformer blocks organized into 4 stages, with window sizes of 8x8 and shifted "
        "window attention for cross-window communication."
    )

    # Add a table for architecture comparison
    doc.add_paragraph("")
    table_caption = doc.add_paragraph()
    table_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tc_run = table_caption.add_run("Table 1: Architecture Comparison Summary")
    tc_run.bold = True
    tc_run.font.size = Pt(10)

    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers_row = ["Architecture", "Parameters (M)", "FLOPS (G)", "RMSE (K)"]
    for i, h in enumerate(headers_row):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    arch_data = [
        ["U-Net Baseline", "23.4", "156.2", "1.84"],
        ["Vision Transformer", "86.7", "312.5", "1.62"],
        ["FourCastNet", "45.2", "198.7", "1.51"],
        ["ClimateFormer", "67.8", "245.3", "1.38"],
    ]
    for r, row_data in enumerate(arch_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph(
        "The temporal decoder employs 6 layers of causal self-attention with 12 attention heads "
        "and a hidden dimension of 768. Cross-attention layers between the spatial encoder output "
        "and the temporal decoder enable dynamic conditioning on the current atmospheric state. "
        "A physics constraint module, implemented as a differentiable finite-difference operator, "
        "projects the decoder output onto the manifold of physically consistent states."
    )

    doc.add_heading("4.3 Graph Neural Network Variant", level=2)
    doc.add_paragraph(
        "As an additional comparison, we implemented a graph neural network (GNN) variant based "
        "on the GraphCast architecture. This model represents the atmosphere as an icosahedral "
        "mesh with approximately 40,000 nodes, connected through a multi-scale graph structure "
        "with edges at three resolution levels. Message passing occurs over 16 processor layers, "
        "each performing neighborhood aggregation with learned edge features that encode distance "
        "and directional relationships between grid points."
    )

    doc.add_page_break()

    # --- Chapter 5: Experimental Results (pages 13-14) ---
    doc.add_heading("5. Experimental Results", level=1)
    doc.add_heading("5.1 Temperature Prediction Performance", level=2)
    doc.add_paragraph(
        "Table 2 presents the temperature prediction performance across all model architectures "
        "at lead times ranging from 24 hours to 10 days. The ClimateFormer achieves the lowest "
        "RMSE at all lead times beyond 48 hours, with particularly notable improvements at the "
        "7-10 day range where traditional models typically exhibit rapid skill degradation."
    )

    table2_cap = doc.add_paragraph()
    table2_cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    t2r = table2_cap.add_run("Table 2: Temperature RMSE (K) by Lead Time")
    t2r.bold = True
    t2r.font.size = Pt(10)

    table2 = doc.add_table(rows=5, cols=5)
    table2.style = "Table Grid"
    t2_headers = ["Model", "24h", "72h", "168h", "240h"]
    for i, h in enumerate(t2_headers):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    t2_data = [
        ["U-Net", "1.12", "1.84", "2.91", "3.67"],
        ["ViT", "1.08", "1.62", "2.54", "3.21"],
        ["GNN", "1.05", "1.53", "2.38", "2.98"],
        ["ClimateFormer", "1.03", "1.38", "2.12", "2.65"],
    ]
    for r, row_data in enumerate(t2_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_heading("5.2 Precipitation Forecasting", level=2)
    doc.add_paragraph(
        "Precipitation prediction remains one of the most challenging aspects of weather "
        "forecasting due to its inherently stochastic and highly localized nature. The CRPS "
        "scores reveal that the ClimateFormer outperforms competing architectures for stratiform "
        "precipitation events but shows more modest improvements for convective rainfall, "
        "suggesting that local-scale processes remain difficult to capture even with advanced "
        "attention mechanisms. The GNN variant shows competitive performance for tropical "
        "precipitation patterns, likely benefiting from its explicit representation of "
        "geographical relationships through the mesh structure."
    )

    doc.add_page_break()

    # --- Chapter 6: Climate Projection Analysis (pages 15-16) ---
    doc.add_heading("6. Climate Projection Analysis", level=1)
    doc.add_heading("6.1 Decadal Projections Under SSP Scenarios", level=2)
    doc.add_paragraph(
        "We evaluate the trained models' ability to generate decadal climate projections under "
        "the Shared Socioeconomic Pathway (SSP) scenarios SSP2-4.5 and SSP5-8.5. Unlike "
        "weather forecasting, climate projections require capturing the statistical properties "
        "of the climate system rather than specific weather events. The ClimateFormer's "
        "physics-informed constraint layer proves particularly valuable in this context, "
        "preventing the model from drifting into unphysical states during extended autoregressive "
        "rollouts."
    )
    doc.add_heading("6.2 Regional Temperature Trends", level=2)
    doc.add_paragraph(
        "Analysis of regional temperature trends reveals that the ClimateFormer projections "
        "align closely with CMIP6 multi-model ensemble means for large-scale patterns but "
        "provide significantly sharper spatial detail. For the European domain, the model "
        "projects warming of 2.3-2.8 K under SSP2-4.5 and 4.1-4.9 K under SSP5-8.5 by 2100, "
        "consistent with the IPCC AR6 assessed ranges. The Arctic amplification factor "
        "predicted by ClimateFormer (2.7) agrees well with observational estimates and exceeds "
        "the CMIP6 ensemble mean (2.4), suggesting the model captures key polar feedback "
        "mechanisms."
    )
    doc.add_heading("6.3 Extreme Event Analysis", level=2)
    doc.add_paragraph(
        "The frequency and intensity of extreme temperature events show marked increases in "
        "all projection scenarios. Under SSP5-8.5, the model projects that current 1-in-50-year "
        "heat events will occur every 5-8 years by mid-century across temperate regions. "
        "Crucially, the ClimateFormer's probabilistic output layer enables direct estimation "
        "of return periods without requiring separate extreme value analysis, a significant "
        "advantage over deterministic model outputs. The model also captures the asymmetric "
        "response of hot and cold extremes, with cold extremes warming faster than hot extremes "
        "in high-latitude regions, consistent with observed trends."
    )

    doc.add_page_break()

    # --- Chapter 7: Discussion (pages 17-18) ---
    doc.add_heading("7. Discussion", level=1)
    doc.add_heading("7.1 Advantages of the Hybrid Approach", level=2)
    doc.add_paragraph(
        "The results demonstrate that the ClimateFormer's hybrid approach, combining data-driven "
        "learning with physical constraints, offers several distinct advantages over purely "
        "empirical or purely physics-based methods. First, the physics constraint layer ensures "
        "global conservation of mass and energy, a property that many purely data-driven models "
        "violate during extended rollouts, leading to unrealistic drift in temperature and moisture "
        "fields. Second, the attention mechanism enables the model to dynamically focus on "
        "regions of active weather development, effectively allocating computational resources "
        "where they are most needed."
    )
    doc.add_heading("7.2 Limitations and Challenges", level=2)
    doc.add_paragraph(
        "Several important limitations merit discussion. The model's reliance on ERA5 reanalysis "
        "for training means that systematic biases in the reanalysis product may be inherited by "
        "the neural network. This is particularly concerning for data-sparse regions such as the "
        "Southern Ocean and African continent, where the reanalysis quality is lower. Additionally, "
        "the model's performance for sub-seasonal to seasonal (S2S) prediction timescales shows "
        "limited improvement over climatological baselines, suggesting that the current architecture "
        "may not adequately capture slow-evolving modes of variability such as the Madden-Julian "
        "Oscillation and stratospheric dynamics."
    )
    doc.add_heading("7.3 Computational Considerations", level=2)
    doc.add_paragraph(
        "From a computational perspective, the ClimateFormer achieves inference speeds approximately "
        "10,000 times faster than equivalent CESM simulations when run on a single A100 GPU. "
        "This dramatic speedup enables ensemble generation at scales previously infeasible, "
        "allowing for comprehensive uncertainty quantification through 1,000-member ensembles "
        "that would require months of supercomputer time using traditional models. However, the "
        "training cost remains substantial, requiring approximately 15,000 GPU-hours for the "
        "full training procedure, underscoring the need for efficient training strategies and "
        "transfer learning approaches."
    )

    doc.add_page_break()

    # --- Chapter 8: Conclusions (page 19) ---
    doc.add_heading("8. Conclusions", level=1)
    doc.add_paragraph(
        "This report has presented a comprehensive investigation into advanced neural network "
        "architectures for climate modeling, culminating in the development of the ClimateFormer "
        "architecture. Our key findings can be summarized as follows:"
    )
    conclusions = [
        "The ClimateFormer achieves state-of-the-art temperature prediction accuracy with RMSE "
        "reductions of 12-28% compared to baseline architectures across all forecast lead times.",
        "Physics-informed constraints are essential for maintaining physical consistency during "
        "extended climate projections, preventing unrealistic drift in conservation quantities.",
        "The hybrid attention-based architecture effectively captures both local convective "
        "processes and global teleconnection patterns, outperforming purely convolutional approaches.",
        "Decadal climate projections under SSP scenarios agree with CMIP6 ensemble ranges while "
        "providing substantially higher spatial resolution and enabling rapid ensemble generation.",
        "Significant challenges remain for sub-seasonal timescales and convective precipitation "
        "prediction, suggesting directions for future architectural innovations.",
    ]
    for item in conclusions:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph(
        "Future work will focus on extending the architecture to coupled ocean-atmosphere "
        "prediction, incorporating satellite observational data for improved initialization, "
        "and developing transfer learning strategies to reduce training computational requirements. "
        "The integration of these neural network approaches with operational forecasting systems "
        "represents a promising pathway toward more accurate and computationally efficient "
        "climate services for society."
    )

    doc.add_page_break()

    # --- References (page 20) ---
    doc.add_heading("References", level=1)
    references = [
        "Beucler, T., Gentine, P., Yuval, J., Gupta, A., & Peng, L. (2024). Climate-invariant machine learning. Science Advances, 10(6), eadj7250.",
        "Bi, K., Xie, L., Zhang, H., Chen, X., Gu, X., & Tian, Q. (2023). Accurate medium-range global weather forecasting with 3D neural networks. Nature, 619, 533-538.",
        "Henderson, G. R., Peings, Y., Furtado, J. C., & Kushner, P. J. (2023). Snow-atmosphere coupling in the Northern Hemisphere. Nature Climate Change, 8, 954-963.",
        "Kim, S., & Park, S. (2024). Physics-constrained deep learning for climate model emulation. Journal of Advances in Modeling Earth Systems, 16(3), e2023MS004012.",
        "Krasnopolsky, V. M., & Fox-Rabinovitz, M. S. (2006). Complex hybrid models combining deterministic and machine learning components. Neural Networks, 19(2), 122-134.",
        "Lam, R., Sanchez-Gonzalez, A., Willson, M., et al. (2023). Learning skillful medium-range global weather forecasting. Science, 382(6677), 1416-1421.",
        "Nguyen, T., Brandstetter, J., Kapoor, A., Gupta, J., & Grover, A. (2022). ClimaX: A foundation model for weather and climate. Proceedings of the 40th ICML, 25904-25938.",
        "Pathak, J., Subramanian, S., Harrington, P., et al. (2022). FourCastNet: A global data-driven high-resolution weather forecasting model. arXiv preprint arXiv:2202.11214.",
        "Ramirez, A. L., & Okonkwo, C. (2024). Evaluation of CMIP6 models over tropical regions. Climate Dynamics, 62, 1847-1865.",
        "Vandal, T., Kodra, E., Ganguly, S., Michaelis, A., Nemani, R., & Ganguly, A. R. (2017). DeepSD: Generating high fidelity daily climate scenarios. Proceedings of the 23rd ACM SIGKDD, 1466-1475.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
