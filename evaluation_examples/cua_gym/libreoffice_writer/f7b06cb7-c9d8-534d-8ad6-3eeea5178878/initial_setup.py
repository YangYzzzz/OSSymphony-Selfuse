"""
Initial Setup: Enable track changes recording in NDA contract document
Task ID: writer_rm_001
Domain: libreoffice_writer

Creates a 3-page NDA Agreement document with track changes DISABLED (default).
Opens it in LibreOffice Writer.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Confidential Business Agreement')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph('')  # spacer

    # --- Parties ---
    doc.add_heading('1. PARTIES', level=1)
    doc.add_paragraph(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of March 15, 2025, '
        'by and between:'
    )
    doc.add_paragraph(
        'Meridian Technologies Inc., a Delaware corporation with its principal offices at '
        '2847 Innovation Drive, Suite 400, San Jose, CA 95134 ("Disclosing Party");',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Pinnacle Consulting Group LLC, a New York limited liability company with its '
        'principal offices at 1200 Park Avenue, 18th Floor, New York, NY 10028 ("Receiving Party").',
        style='List Bullet'
    )

    # --- Purpose ---
    doc.add_heading('2. PURPOSE', level=1)
    doc.add_paragraph(
        'The Disclosing Party possesses certain confidential and proprietary information '
        'relating to its quantum computing research platform, codenamed "Project Aurora," '
        'including but not limited to technical specifications, business strategies, financial '
        'projections, customer data, and product roadmaps. The Disclosing Party wishes to '
        'share such information with the Receiving Party for the purpose of evaluating a '
        'potential strategic partnership and joint venture opportunity in the Asia-Pacific '
        'market, specifically targeting enterprise clients in Japan, South Korea, and Australia.'
    )

    # --- Definition of Confidential Information ---
    doc.add_heading('3. DEFINITION OF CONFIDENTIAL INFORMATION', level=1)
    doc.add_paragraph(
        '"Confidential Information" shall mean any and all non-public information, whether '
        'oral, written, electronic, or visual, that is disclosed by the Disclosing Party to '
        'the Receiving Party, including but not limited to:'
    )
    items = [
        'Technical data, trade secrets, know-how, research findings, inventions, processes, '
        'techniques, algorithms, software programs (including source code and object code), '
        'databases, and hardware configurations;',
        'Business information including financial statements, projections, budgets, marketing '
        'plans, customer lists, supplier agreements, pricing structures, and sales data;',
        'Product designs, prototypes, specifications, engineering drawings, manufacturing '
        'processes, and quality control procedures;',
        'Strategic plans, merger and acquisition targets, partnership proposals, and '
        'organizational restructuring plans;',
        'Employee information, compensation structures, personnel records, and recruitment '
        'strategies;',
        'Any information that is marked as "Confidential," "Proprietary," or with a similar '
        'designation, or that a reasonable person would understand to be confidential given '
        'the nature of the information and circumstances of disclosure.',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # --- Obligations ---
    doc.add_heading('4. OBLIGATIONS OF THE RECEIVING PARTY', level=1)
    doc.add_paragraph(
        'The Receiving Party agrees to hold and maintain the Confidential Information in '
        'strict confidence for the sole and exclusive benefit of the Disclosing Party. '
        'Specifically, the Receiving Party shall:'
    )
    obligations = [
        'Not disclose, publish, or otherwise reveal any Confidential Information to any '
        'third party without the prior written consent of the Disclosing Party;',
        'Use the Confidential Information solely for the purpose of evaluating and pursuing '
        'the potential strategic partnership described in Section 2 above;',
        'Limit access to the Confidential Information to those employees, agents, and '
        'advisors who have a demonstrable need to know and who have executed binding '
        'confidentiality agreements no less restrictive than this Agreement;',
        'Take all reasonable precautions to prevent unauthorized disclosure or use of the '
        'Confidential Information, including implementing appropriate physical, electronic, '
        'and procedural safeguards;',
        'Promptly notify the Disclosing Party in writing upon discovery of any unauthorized '
        'use or disclosure of Confidential Information, and cooperate fully with the '
        'Disclosing Party in remedying such breach.',
    ]
    for ob in obligations:
        doc.add_paragraph(ob, style='List Number')

    # --- Exclusions --- (page 2 territory)
    doc.add_heading('5. EXCLUSIONS FROM CONFIDENTIAL INFORMATION', level=1)
    doc.add_paragraph(
        'The obligations set forth in Section 4 shall not apply to information that:'
    )
    exclusions = [
        'Was publicly known and generally available in the public domain prior to the time '
        'of disclosure by the Disclosing Party;',
        'Becomes publicly known and generally available after disclosure by the Disclosing '
        'Party through no wrongful act, fault, or negligence of the Receiving Party;',
        'Was already in the lawful possession of the Receiving Party at the time of '
        'disclosure, as evidenced by contemporaneous written records;',
        'Is independently developed by the Receiving Party without use of or reference to '
        'the Confidential Information, as demonstrated by documented evidence;',
        'Is obtained by the Receiving Party from a third party without breach of any '
        'obligation of confidentiality.',
    ]
    for ex in exclusions:
        doc.add_paragraph(ex, style='List Number')

    # --- Term ---
    doc.add_heading('6. TERM AND TERMINATION', level=1)
    doc.add_paragraph(
        'This Agreement shall remain in effect for a period of three (3) years from the '
        'Effective Date stated above, unless terminated earlier by either party upon thirty '
        '(30) days\' prior written notice to the other party. The obligations of '
        'confidentiality contained herein shall survive any termination or expiration of '
        'this Agreement for an additional period of five (5) years following the date of '
        'termination or expiration.'
    )
    doc.add_paragraph(
        'Upon termination or expiration of this Agreement, the Receiving Party shall '
        'promptly return or destroy all documents, materials, and other tangible '
        'manifestations of Confidential Information in its possession, including all copies, '
        'extracts, and summaries thereof. The Receiving Party shall certify in writing to '
        'the Disclosing Party that all such materials have been returned or destroyed within '
        'fifteen (15) business days of the termination or expiration date.'
    )

    # --- Remedies ---
    doc.add_heading('7. REMEDIES', level=1)
    doc.add_paragraph(
        'The Receiving Party acknowledges and agrees that any breach or threatened breach '
        'of this Agreement may cause irreparable harm to the Disclosing Party, for which '
        'monetary damages alone would be an inadequate remedy. Accordingly, in addition to '
        'any other remedies available at law or in equity, the Disclosing Party shall be '
        'entitled to seek injunctive relief, specific performance, or other equitable '
        'remedies without the necessity of proving actual damages or posting any bond or '
        'other security.'
    )
    doc.add_paragraph(
        'The prevailing party in any legal action arising out of or relating to this '
        'Agreement shall be entitled to recover its reasonable attorneys\' fees, court costs, '
        'and other litigation expenses from the non-prevailing party.'
    )

    # --- Miscellaneous --- (page 3 territory)
    doc.add_heading('8. MISCELLANEOUS PROVISIONS', level=1)

    doc.add_heading('8.1 Governing Law', level=2)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of '
        'the State of Delaware, without regard to its conflict of laws principles. Any '
        'disputes arising under or in connection with this Agreement shall be subject to the '
        'exclusive jurisdiction of the state and federal courts located in Wilmington, Delaware.'
    )

    doc.add_heading('8.2 Entire Agreement', level=2)
    doc.add_paragraph(
        'This Agreement constitutes the entire agreement between the parties with respect to '
        'the subject matter hereof and supersedes all prior negotiations, representations, '
        'warranties, commitments, offers, contracts, and writings, whether written or oral, '
        'with respect thereto.'
    )

    doc.add_heading('8.3 Amendment', level=2)
    doc.add_paragraph(
        'No amendment, modification, or waiver of any provision of this Agreement shall be '
        'effective unless set forth in a written instrument signed by both parties. No waiver '
        'of any breach or default under this Agreement shall constitute a waiver of any '
        'subsequent breach or default.'
    )

    doc.add_heading('8.4 Assignment', level=2)
    doc.add_paragraph(
        'Neither party may assign or transfer this Agreement or any of its rights or '
        'obligations hereunder without the prior written consent of the other party, except '
        'that either party may assign this Agreement to a successor in connection with a '
        'merger, acquisition, or sale of all or substantially all of its assets.'
    )

    doc.add_heading('8.5 Severability', level=2)
    doc.add_paragraph(
        'If any provision of this Agreement is held to be invalid, illegal, or unenforceable '
        'under applicable law, such provision shall be modified to the minimum extent necessary '
        'to make it valid, legal, and enforceable, and the remaining provisions of this '
        'Agreement shall continue in full force and effect.'
    )

    doc.add_heading('8.6 Notices', level=2)
    doc.add_paragraph(
        'All notices, requests, demands, and other communications under this Agreement shall '
        'be in writing and shall be deemed to have been duly given when delivered personally, '
        'sent by certified mail (return receipt requested), or sent by overnight courier to '
        'the addresses set forth in Section 1 above, or to such other address as either party '
        'may designate by written notice to the other party.'
    )

    # --- Signature block ---
    doc.add_paragraph('')  # spacer
    doc.add_heading('SIGNATURES', level=1)
    doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties have executed this Non-Disclosure Agreement as of '
        'the Effective Date first written above.'
    )
    doc.add_paragraph('')

    # Disclosing Party signature
    p1 = doc.add_paragraph()
    p1.add_run('DISCLOSING PARTY:').bold = True
    doc.add_paragraph('Meridian Technologies Inc.')
    doc.add_paragraph('')
    sig1 = doc.add_paragraph('_______________________________')
    doc.add_paragraph('Name: Dr. Elena Vasquez, Chief Executive Officer')
    doc.add_paragraph('Date: _______________')
    doc.add_paragraph('')

    # Receiving Party signature
    p2 = doc.add_paragraph()
    p2.add_run('RECEIVING PARTY:').bold = True
    doc.add_paragraph('Pinnacle Consulting Group LLC')
    doc.add_paragraph('')
    sig2 = doc.add_paragraph('_______________________________')
    doc.add_paragraph('Name: Jonathan R. Whitfield, Managing Partner')
    doc.add_paragraph('Date: _______________')

    # --- Ensure track changes is DISABLED (default, but be explicit) ---
    # Remove any trackChanges element from settings if present
    settings = doc.settings.element
    for tc in settings.findall(qn('w:trackChanges')):
        settings.remove(tc)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
