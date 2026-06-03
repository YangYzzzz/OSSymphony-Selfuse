"""
Reward Script: Recipe card layout with decorative border, colored title banner,
               ingredients/instructions columns in a borderless table.
Task ID: writer_rd_086
Domain: libreoffice_writer
Scoring:
  C1 (0.20) - Page borders: 2pt solid dark red (#8B0000) on all 4 sides
  C2 (0.25) - Title banner: "Classic Margherita Pizza", centered, red background, white bold 22pt
  C3 (0.25) - 2-column borderless table with Ingredients (8 items) and Instructions (6 steps)
  C4 (0.15) - Column headers bold 14pt dark red; content present
  C5 (0.15) - Bottom line: "Prep: 20 min | Cook: 12 min | Serves: 4" centered italic
"""

import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_086'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # ---------------------------------------------------------------
    # Component 1: Page borders - 2pt solid dark red on all 4 sides (0.20 pts)
    # ---------------------------------------------------------------
    try:
        section = doc.sections[0]
        sectPr = section._sectPr
        pgBorders = sectPr.find(qn('w:pgBorders'))
        if pgBorders is not None:
            sides = ['top', 'left', 'bottom', 'right']
            borders_ok = 0
            for side in sides:
                el = pgBorders.find(qn('w:' + side))
                if el is not None:
                    val = el.get(qn('w:val'))
                    sz = el.get(qn('w:sz'))
                    color = el.get(qn('w:color'))
                    # sz=16 means 2pt (in eighth-points), val=single, color=8B0000
                    if val == 'single' and color and color.upper() == '8B0000':
                        borders_ok += 1
                        print("PASS: Page border %s — single, color 8B0000" % side)
                    else:
                        print("FAIL: Page border %s — val=%s, color=%s (expected single, 8B0000)" % (side, val, color))
                else:
                    print("FAIL: Page border %s not found" % side)
            if borders_ok == 4:
                total_score += 0.20
                print("PASS: Component 1 — All 4 page borders correct (0.20 pts)")
            elif borders_ok >= 2:
                total_score += 0.10
                print("PARTIAL: Component 1 — %d/4 borders correct (0.10 pts)" % borders_ok)
            else:
                print("FAIL: Component 1 — Only %d/4 borders correct" % borders_ok)
        else:
            print("FAIL: Component 1 — No page borders found")
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # ---------------------------------------------------------------
    # Component 2: Title banner paragraph (0.25 pts)
    #   - Text "Classic Margherita Pizza"
    #   - Centered alignment
    #   - Paragraph shading fill CC0000 (red background)
    #   - White bold ~22pt text
    # ---------------------------------------------------------------
    try:
        title_found = False
        for para in doc.paragraphs:
            if 'Classic Margherita Pizza' in para.text:
                title_found = True
                c2_sub = 0.0

                # Check centered
                pf_align = para.paragraph_format.alignment
                if pf_align is not None and pf_align == 1:  # CENTER
                    c2_sub += 0.05
                    print("PASS: Title is centered")
                else:
                    # Also check XML jc
                    pPr = para._element.find(qn('w:pPr'))
                    jc = pPr.find(qn('w:jc')) if pPr is not None else None
                    if jc is not None and jc.get(qn('w:val')) == 'center':
                        c2_sub += 0.05
                        print("PASS: Title is centered (XML)")
                    else:
                        print("FAIL: Title not centered, align=%s" % pf_align)

                # Check background shading (CC0000)
                pPr = para._element.find(qn('w:pPr'))
                shd = pPr.find(qn('w:shd')) if pPr is not None else None
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill.upper() in ('CC0000', 'CC0000'):
                        c2_sub += 0.08
                        print("PASS: Title has red background fill CC0000")
                    else:
                        print("FAIL: Title shading fill=%s, expected CC0000" % fill)
                else:
                    print("FAIL: Title has no paragraph shading")

                # Check run properties: white, bold, ~22pt
                if para.runs:
                    run = para.runs[0]
                    # Bold
                    if run.font.bold:
                        c2_sub += 0.04
                        print("PASS: Title is bold")
                    else:
                        print("FAIL: Title not bold")

                    # White color
                    if run.font.color.rgb and str(run.font.color.rgb).upper() == 'FFFFFF':
                        c2_sub += 0.04
                        print("PASS: Title text is white")
                    else:
                        print("FAIL: Title color=%s, expected FFFFFF" % (run.font.color.rgb,))

                    # Size ~22pt (allow 20-24pt)
                    if run.font.size:
                        pt_val = run.font.size.pt
                        if 20 <= pt_val <= 24:
                            c2_sub += 0.04
                            print("PASS: Title size=%.1fpt" % pt_val)
                        else:
                            print("FAIL: Title size=%.1fpt, expected ~22pt" % pt_val)
                    else:
                        print("FAIL: Title has no font size set")
                else:
                    print("FAIL: Title paragraph has no runs")

                total_score += c2_sub
                print("Component 2 subtotal: %.2f/0.25 pts" % c2_sub)
                break

        if not title_found:
            print("FAIL: Component 2 — 'Classic Margherita Pizza' title not found")
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # ---------------------------------------------------------------
    # Component 3: 2-column table with no visible borders,
    #   ingredients (8 items) and instructions (6 steps) (0.25 pts)
    # ---------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No tables found")
        else:
            table = doc.tables[0]
            c3_sub = 0.0

            # Check 2 columns
            if len(table.columns) == 2:
                c3_sub += 0.05
                print("PASS: Table has 2 columns")
            else:
                print("FAIL: Table has %d columns, expected 2" % len(table.columns))

            # Check table borders are 'none' (no visible borders)
            tblPr = table._tbl.find(qn('w:tblPr'))
            if tblPr is not None:
                borders = tblPr.find(qn('w:tblBorders'))
                if borders is not None:
                    all_none = True
                    for child in borders:
                        bval = child.get(qn('w:val'))
                        if bval not in ('none', 'nil'):
                            all_none = False
                    if all_none:
                        c3_sub += 0.05
                        print("PASS: Table borders are invisible")
                    else:
                        print("FAIL: Some table borders are visible")
                else:
                    # No explicit borders element could mean default visible borders
                    print("INFO: No tblBorders element — borders may use style defaults")

            # Check left cell has ingredients (at least 6 bullet items)
            left_cell = table.cell(0, 0)
            ingredient_lines = [p.text.strip() for p in left_cell.paragraphs if p.text.strip().startswith(('•', '-', '*'))]
            if len(ingredient_lines) >= 6:
                c3_sub += 0.05
                print("PASS: Left cell has %d ingredient items (>= 6)" % len(ingredient_lines))
            else:
                print("FAIL: Left cell has %d ingredient items, expected >= 6" % len(ingredient_lines))

            # Check right cell has instructions (numbered steps)
            right_cell = table.cell(0, 1)
            import re
            step_lines = [p.text.strip() for p in right_cell.paragraphs if re.match(r'^\d+\.', p.text.strip())]
            if len(step_lines) >= 4:
                c3_sub += 0.05
                print("PASS: Right cell has %d instruction steps (>= 4)" % len(step_lines))
            else:
                print("FAIL: Right cell has %d instruction steps, expected >= 4" % len(step_lines))

            # Check "Ingredients" header in left cell
            left_first = left_cell.paragraphs[0].text.strip() if left_cell.paragraphs else ''
            right_first = right_cell.paragraphs[0].text.strip() if right_cell.paragraphs else ''
            if 'ingredient' in left_first.lower():
                c3_sub += 0.025
                print("PASS: Left column header contains 'Ingredients'")
            else:
                print("FAIL: Left column header=%s, expected 'Ingredients'" % repr(left_first))

            if 'instruction' in right_first.lower():
                c3_sub += 0.025
                print("PASS: Right column header contains 'Instructions'")
            else:
                print("FAIL: Right column header=%s, expected 'Instructions'" % repr(right_first))

            total_score += c3_sub
            print("Component 3 subtotal: %.3f/0.25 pts" % c3_sub)
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    # ---------------------------------------------------------------
    # Component 4: Column headers bold 14pt dark red (0.15 pts)
    # ---------------------------------------------------------------
    try:
        if len(doc.tables) > 0:
            table = doc.tables[0]
            c4_sub = 0.0

            for ci, expected_header in enumerate(['Ingredients', 'Instructions']):
                cell = table.cell(0, ci)
                if cell.paragraphs:
                    header_para = cell.paragraphs[0]
                    header_runs = header_para.runs
                    if header_runs:
                        hr = header_runs[0]
                        checks = 0
                        # Bold
                        if hr.font.bold:
                            checks += 1
                        # Size ~14pt
                        if hr.font.size and 12 <= hr.font.size.pt <= 16:
                            checks += 1
                        # Dark red color
                        if hr.font.color.rgb and str(hr.font.color.rgb).upper() == '8B0000':
                            checks += 1

                        if checks == 3:
                            c4_sub += 0.075
                            print("PASS: '%s' header — bold, ~14pt, dark red" % expected_header)
                        elif checks >= 2:
                            c4_sub += 0.05
                            print("PARTIAL: '%s' header — %d/3 properties correct" % (expected_header, checks))
                        else:
                            print("FAIL: '%s' header — only %d/3 properties (bold=%s, size=%s, color=%s)" % (
                                expected_header, checks, hr.font.bold, hr.font.size, hr.font.color.rgb))
                    else:
                        print("FAIL: '%s' header has no runs" % expected_header)

            total_score += c4_sub
            print("Component 4 subtotal: %.3f/0.15 pts" % c4_sub)
        else:
            print("FAIL: Component 4 — No tables found")
    except Exception as e:
        print("ERROR: Component 4 — %s" % e)

    # ---------------------------------------------------------------
    # Component 5: Bottom line centered italic with prep/cook/serves info (0.15 pts)
    # ---------------------------------------------------------------
    try:
        c5_sub = 0.0
        prep_para = None
        for para in doc.paragraphs:
            txt = para.text.strip()
            if 'prep' in txt.lower() and 'cook' in txt.lower() and 'serves' in txt.lower():
                prep_para = para
                break

        if prep_para is not None:
            c5_sub += 0.05
            print("PASS: Found prep/cook/serves paragraph")

            # Check centered
            pf_align = prep_para.paragraph_format.alignment
            if pf_align is not None and pf_align == 1:
                c5_sub += 0.05
                print("PASS: Prep line is centered")
            else:
                pPr = prep_para._element.find(qn('w:pPr'))
                jc = pPr.find(qn('w:jc')) if pPr is not None else None
                if jc is not None and jc.get(qn('w:val')) == 'center':
                    c5_sub += 0.05
                    print("PASS: Prep line is centered (XML)")
                else:
                    print("FAIL: Prep line not centered, align=%s" % pf_align)

            # Check italic
            if prep_para.runs:
                if prep_para.runs[0].font.italic:
                    c5_sub += 0.05
                    print("PASS: Prep line is italic")
                else:
                    print("FAIL: Prep line not italic")
            else:
                print("FAIL: Prep paragraph has no runs")
        else:
            print("FAIL: Component 5 — No paragraph with prep/cook/serves info found")

        total_score += c5_sub
        print("Component 5 subtotal: %.2f/0.15 pts" % c5_sub)
    except Exception as e:
        print("ERROR: Component 5 — %s" % e)

    final_score = min(round(total_score, 2), 1.0)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print("PERSIST_WARN: save hook failed: %s" % e)


# Entry point
file_path = os.path.join(WORKDIR, TASK_ID + '.docx')
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
