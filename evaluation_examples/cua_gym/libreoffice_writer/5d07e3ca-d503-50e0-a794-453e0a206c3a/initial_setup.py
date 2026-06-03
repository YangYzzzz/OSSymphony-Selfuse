"""
Initial Setup: Create a Writer document with code snippets in Default Paragraph Style.
Task ID: writer_bs_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Python Data Processing Guide", level=1)

    # Introductory paragraph
    doc.add_paragraph(
        "This guide covers essential Python techniques for processing "
        "and analyzing datasets. Each section includes code examples "
        "that demonstrate practical usage patterns for common data operations."
    )

    # Section 1
    doc.add_heading("1. Reading CSV Files", level=2)
    doc.add_paragraph(
        "The pandas library provides a simple interface for reading structured "
        "data from CSV files. Below is a basic example of loading a dataset:"
    )

    # Code snippet in default style (no special formatting)
    doc.add_paragraph(
        'import pandas as pd\n'
        '\n'
        'df = pd.read_csv("sales_report_2025.csv")\n'
        'print(f"Loaded {len(df)} records")\n'
        'print(df.head())'
    )

    doc.add_paragraph(
        "After loading the data, you can inspect the first few rows using the "
        "head() method to verify the structure matches your expectations."
    )

    # Section 2
    doc.add_heading("2. Filtering and Aggregation", level=2)
    doc.add_paragraph(
        "Once the data is loaded, filtering rows based on conditions and "
        "computing aggregate statistics are among the most frequent operations."
    )

    doc.add_paragraph(
        'q3_sales = df[df["quarter"] == "Q3"]\n'
        'total_revenue = q3_sales["revenue"].sum()\n'
        'avg_order = q3_sales["order_value"].mean()\n'
        'print(f"Q3 Revenue: ${total_revenue:,.2f}")\n'
        'print(f"Average Order: ${avg_order:,.2f}")'
    )

    doc.add_paragraph(
        "The filtering syntax uses boolean indexing, which is both readable "
        "and efficient for large datasets."
    )

    # Section 3
    doc.add_heading("3. Data Visualization", level=2)
    doc.add_paragraph(
        "Matplotlib integrates well with pandas DataFrames, allowing you to "
        "create publication-quality charts with minimal code."
    )

    doc.add_paragraph(
        'import matplotlib.pyplot as plt\n'
        '\n'
        'fig, ax = plt.subplots(figsize=(10, 6))\n'
        'monthly = df.groupby("month")["revenue"].sum()\n'
        'monthly.plot(kind="bar", ax=ax, color="#4472C4")\n'
        'ax.set_title("Monthly Revenue Overview")\n'
        'ax.set_ylabel("Revenue ($)")\n'
        'plt.tight_layout()\n'
        'plt.savefig("revenue_chart.png", dpi=150)'
    )

    doc.add_paragraph(
        "Always call tight_layout() before saving to prevent label clipping, "
        "especially when using rotated axis labels or multi-line titles."
    )

    # Section 4
    doc.add_heading("4. Error Handling in Data Pipelines", level=2)
    doc.add_paragraph(
        "Production data pipelines should anticipate missing or malformed "
        "data. Using try-except blocks with specific exception types ensures "
        "graceful degradation."
    )

    doc.add_paragraph(
        'def process_record(record):\n'
        '    try:\n'
        '        value = float(record["amount"])\n'
        '        if value < 0:\n'
        '            raise ValueError("Negative amount")\n'
        '        return value * 1.08  # Apply tax\n'
        '    except (KeyError, TypeError) as e:\n'
        '        logging.warning(f"Skipped record: {e}")\n'
        '        return None'
    )

    doc.add_paragraph(
        "This approach logs warnings for problematic records while allowing "
        "the pipeline to continue processing the remaining data."
    )

    # Section 5
    doc.add_heading("5. Exporting Results", level=2)
    doc.add_paragraph(
        "After processing, results typically need to be written to files "
        "or databases for downstream consumption."
    )

    doc.add_paragraph(
        'output_df = df[["product_id", "name", "adjusted_price"]]\n'
        'output_df.to_excel("processed_output.xlsx", index=False)\n'
        'output_df.to_json("processed_output.json", orient="records")\n'
        'print(f"Exported {len(output_df)} records")'
    )

    doc.add_paragraph(
        "When exporting to Excel, set index=False to avoid writing the "
        "DataFrame index as an extra column, which can confuse downstream "
        "consumers expecting a clean tabular structure."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
