"""
Initial Setup: Benefits Enrollment Guide - no comparison table yet
Task ID: writer_hr_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Benefits Enrollment Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Fiscal Year 2025-2026 Open Enrollment')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()  # spacer

    # --- Section 1: Overview ---
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'Welcome to the Annual Benefits Enrollment period for Meridian Technologies. '
        'This guide provides essential information about the health insurance plans '
        'available to all eligible employees and their dependents. During the open '
        'enrollment window, you may elect new coverage, change your current plan, '
        'add or remove dependents, or waive coverage entirely.'
    )
    doc.add_paragraph(
        'Meridian Technologies is committed to offering competitive benefits that '
        'support the health and well-being of our workforce. This year, we continue '
        'to partner with BlueCross BlueShield to provide three distinct health '
        'insurance tiers designed to meet a range of healthcare needs and budgets.'
    )

    # --- Section 2: Eligibility ---
    doc.add_heading('Eligibility', level=1)
    doc.add_paragraph(
        'All full-time employees who have completed their 90-day probationary period '
        'are eligible for benefits enrollment. Part-time employees working 30 or more '
        'hours per week are also eligible for coverage. Eligible dependents include '
        'legal spouses, domestic partners, and children under age 26.'
    )
    doc.add_paragraph(
        'New hires must enroll within 30 calendar days of their start date. If you '
        'miss the enrollment window, you will need to wait until the next open '
        'enrollment period unless you experience a qualifying life event such as '
        'marriage, birth of a child, or loss of other coverage.'
    )

    # --- Section 3: Enrollment Timeline ---
    doc.add_heading('Enrollment Timeline', level=1)
    doc.add_paragraph('Please note the following key dates for the upcoming enrollment cycle:')
    doc.add_paragraph('Open Enrollment Begins: November 1, 2025', style='List Bullet')
    doc.add_paragraph('Open Enrollment Ends: November 30, 2025', style='List Bullet')
    doc.add_paragraph('Coverage Effective Date: January 1, 2026', style='List Bullet')
    doc.add_paragraph('Enrollment Confirmation Letters Mailed: December 15, 2025', style='List Bullet')

    # --- Section 4: Available Plans ---
    doc.add_heading('Available Health Insurance Plans', level=1)
    doc.add_paragraph(
        'Meridian Technologies offers three health insurance plan tiers through '
        'BlueCross BlueShield. Each plan differs in monthly premium costs, deductible '
        'amounts, and out-of-pocket maximums. Employees are encouraged to carefully '
        'review the plan details before making their selection.'
    )
    doc.add_paragraph(
        'A detailed comparison of the available plans will help you determine which '
        'option best fits your healthcare needs and financial situation. Please review '
        'the plan options below and consult with Human Resources if you have questions.'
    )

    # NOTE: No table here — the task is for the agent to add the comparison table.

    # --- Section 5: How to Enroll ---
    doc.add_heading('How to Enroll', level=1)
    doc.add_paragraph(
        'To complete your enrollment, log in to the Meridian HR Portal at '
        'https://hr.meridiantech.com and navigate to Benefits > Open Enrollment. '
        'Follow the on-screen instructions to review plan options, add dependents, '
        'and confirm your selections.'
    )
    doc.add_paragraph(
        'If you encounter technical difficulties, contact the IT Help Desk at '
        'ext. 4500 or email helpdesk@meridiantech.com. For benefits-specific '
        'questions, reach out to the HR Benefits Team at benefits@meridiantech.com '
        'or call ext. 3200.'
    )

    # --- Section 6: Additional Resources ---
    doc.add_heading('Additional Resources', level=1)
    doc.add_paragraph('Employee Assistance Program (EAP): 1-800-555-0199', style='List Bullet')
    doc.add_paragraph('BlueCross BlueShield Member Services: 1-800-555-0234', style='List Bullet')
    doc.add_paragraph('HR Benefits Team Office Hours: Monday-Friday, 9:00 AM - 5:00 PM', style='List Bullet')
    doc.add_paragraph('Benefits FAQ: https://hr.meridiantech.com/benefits-faq', style='List Bullet')

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = closing.add_run('Thank you for being part of the Meridian Technologies team!')
    run.font.italic = True
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
