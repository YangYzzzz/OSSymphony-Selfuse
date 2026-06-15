"""
Reward Script: Add health insurance plans comparison table with alternating row colors
Task ID: writer_hr_037
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.3): Table exists with correct structure (4 rows x 5 cols, correct headers)
  - Component 2 (0.4): Table contains correct plan data for all 3 plans
  - Component 3 (0.3): Alternating row colors are applied to the table rows
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_037'

# Expected headers
EXPECTED_HEADERS = ['Plan Name', 'Monthly Premium (Employee)', 'Monthly Premium (Family)', 'Deductible', 'Out-of-Pocket Max']

# Expected plan data (plan_name -> [employee_premium, family_premium, deductible, oop_max])
# We normalize by stripping $ and commas for comparison
EXPECTED_PLANS = {
    'basic':    ['150', '400', '2000', '6000'],
    'standard': ['280', '700', '1000', '4000'],
    'premium':  ['450', '1100', '500', '2000'],
}


def normalize_currency(val):
    """Strip $, commas, spaces, and leading zeros from a currency string."""
    if val is None:
        return ''
    return val.strip().replace('$', '').replace(',', '').replace(' ', '').lstrip('0') or '0'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table structure — 4 rows x 5 cols with correct headers (0.3 points)
    try:
        structure_ok = True
        # Check dimensions
        if num_rows < 4:
            print(f"FAIL: Component 1 — Expected at least 4 rows, found {num_rows}")
            structure_ok = False
        if num_cols != 5:
            print(f"FAIL: Component 1 — Expected 5 columns, found {num_cols}")
            structure_ok = False

        if structure_ok:
            # Check header row
            actual_headers = [table.cell(0, c).text.strip() for c in range(5)]
            headers_match = True
            for i, (exp, act) in enumerate(zip(EXPECTED_HEADERS, actual_headers)):
                if exp.lower() != act.lower():
                    print(f"FAIL: Component 1 — Header col {i}: expected '{exp}', found '{act}'")
                    headers_match = False

            if headers_match:
                print(f"PASS: Component 1 — Table structure correct: {num_rows} rows x {num_cols} cols, headers match (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — Headers do not match expected values")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Plan data correctness (0.4 points)
    # Each plan is worth ~0.133 points; partial credit per plan
    try:
        plans_found = 0
        plans_correct = 0

        for row_idx in range(1, min(num_rows, 4)):
            plan_name = table.cell(row_idx, 0).text.strip().lower()
            if plan_name in EXPECTED_PLANS:
                plans_found += 1
                expected_vals = EXPECTED_PLANS[plan_name]
                actual_vals = [normalize_currency(table.cell(row_idx, c).text) for c in range(1, 5)]

                if actual_vals == expected_vals:
                    plans_correct += 1
                    print(f"PASS: Component 2 — Plan '{plan_name}' data correct")
                else:
                    print(f"FAIL: Component 2 — Plan '{plan_name}': expected {expected_vals}, found {actual_vals}")
            else:
                print(f"FAIL: Component 2 — Row {row_idx} plan name '{plan_name}' not recognized")

        if plans_correct == 3:
            print(f"PASS: Component 2 — All 3 plans have correct data (0.4 pts)")
            total_score += 0.4
        elif plans_correct > 0:
            partial = round(0.4 * plans_correct / 3, 2)
            print(f"PARTIAL: Component 2 — {plans_correct}/3 plans correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No plans have correct data")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Alternating row colors (0.3 points)
    # Rows must have at least 2 distinct fill colors to count as alternating
    try:
        row_fills = []
        for ri in range(num_rows):
            # Get fill color of first cell in each row
            tc = table.rows[ri].cells[0]._tc
            tcPr = tc.find(qn('w:tcPr'))
            fill = None
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
            row_fills.append(fill)

        # Check that we have shading info and at least 2 distinct colors
        non_none_fills = [f for f in row_fills if f is not None]

        if len(non_none_fills) < 2:
            print(f"FAIL: Component 3 — Insufficient shading data: {row_fills}")
        else:
            distinct_fills = set(non_none_fills)
            if len(distinct_fills) >= 2:
                # Verify alternating pattern: adjacent rows should differ
                alternating = False
                for i in range(len(row_fills) - 1):
                    if row_fills[i] is not None and row_fills[i + 1] is not None:
                        if row_fills[i] != row_fills[i + 1]:
                            alternating = True
                            break

                if alternating:
                    print(f"PASS: Component 3 — Alternating row colors detected: {row_fills} (0.3 pts)")
                    total_score += 0.3
                else:
                    print(f"FAIL: Component 3 — Rows have colors but not alternating: {row_fills}")
            else:
                print(f"FAIL: Component 3 — All rows have the same fill color: {distinct_fills}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
