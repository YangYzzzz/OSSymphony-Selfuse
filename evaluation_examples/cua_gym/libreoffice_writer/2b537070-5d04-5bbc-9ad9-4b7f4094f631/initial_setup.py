"""
Initial Setup: Legal document with manually numbered clauses (no auto-numbering)
Task ID: writer_pd_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('SERVICE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        'This Service Agreement ("Agreement") is entered into as of March 15, 2025, '
        'by and between Meridian Technology Solutions Inc. ("Provider") and '
        'Crestwood Financial Group LLC ("Client"). The parties agree to the following '
        'terms and conditions governing the provision of IT consulting and managed '
        'services as described herein.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # 15 clause paragraphs with MANUAL numbering (inconsistent indentation)
    # Structure:
    # 1. Scope of Services (L1)
    #   1.1 Provider shall deliver... (L2)
    #   1.2 All deliverables shall... (L2)
    #     1.2.1 Software modules... (L3)
    #     1.2.2 Documentation packages... (L3)
    # 2. Term and Termination (L1)
    #   2.1 This Agreement commences... (L2)
    #   2.2 Either party may terminate... (L2)
    #     2.2.1 Written notice of no less... (L3)
    #     2.2.2 Material breach that remains... (L3)
    # 3. Compensation and Payment (L1)
    #   3.1 Client shall pay Provider... (L2)
    #     3.1.1 Monthly retainer fees... (L3)
    #     3.1.2 Additional project-based... (L3)
    #   3.2 Payment terms are net... (L2)
    # 4. Confidentiality (L1)
    #   4.1 Both parties agree to maintain... (L2)

    clauses = [
        # (text, level, manual_number, indent_inches - deliberately inconsistent)
        ("1. Scope of Services. Provider agrees to furnish Client with comprehensive IT consulting services, including but not limited to network architecture design, cybersecurity assessment, cloud migration planning, and ongoing technical support as outlined in Exhibit A attached hereto.", 1, 0),
        ("1.1 Provider shall deliver all services in accordance with industry best practices and applicable regulatory standards, including SOC 2 Type II compliance requirements and ISO 27001 security frameworks.", 2, 0.4),
        ("1.2 All deliverables produced under this Agreement shall conform to the specifications set forth in the Statement of Work and shall include the following components:", 2, 0.6),
        ("1.2.1 Software modules and integrations developed specifically for Client's enterprise resource planning system, tested against the acceptance criteria defined in Section 7 of the Statement of Work.", 3, 0.8),
        ("1.2.2 Documentation packages including technical specifications, user manuals, administrator guides, and training materials sufficient for Client's internal IT team to maintain and operate the delivered systems.", 3, 1.1),
        ("2. Term and Termination. This Agreement shall govern the relationship between the parties for the duration specified herein and may be terminated only in accordance with the provisions of this Section.", 1, 0),
        ("2.1 This Agreement commences on March 15, 2025 and shall remain in effect for an initial period of twenty-four (24) months, automatically renewing for successive twelve-month periods unless terminated pursuant to subsection 2.2 below.", 2, 0.5),
        ("2.2 Either party may terminate this Agreement prior to the expiration of the then-current term upon the occurrence of any of the following conditions:", 2, 0.45),
        ("2.2.1 Written notice of no less than ninety (90) calendar days delivered to the other party via certified mail or nationally recognized overnight courier service to the address specified in Section 12 of this Agreement.", 3, 1.0),
        ("2.2.2 Material breach that remains uncured for a period of thirty (30) business days following written notification specifying the nature of the breach and the actions required to cure such breach.", 3, 0.9),
        ("3. Compensation and Payment. Client acknowledges its obligation to compensate Provider for all services rendered and expenses incurred in connection with the performance of this Agreement.", 1, 0),
        ("3.1 Client shall pay Provider in accordance with the fee schedule attached hereto as Exhibit B, which includes the following categories of charges:", 2, 0.55),
        ("3.1.1 Monthly retainer fees of Forty-Five Thousand Dollars ($45,000) due on the first business day of each calendar month, covering baseline managed services and up to one hundred sixty (160) hours of consulting time.", 3, 1.05),
        ("3.1.2 Additional project-based fees calculated at the rate of Three Hundred Twenty-Five Dollars ($325) per hour for work exceeding the monthly retainer allocation, invoiced bi-weekly with detailed time records.", 3, 0.95),
        ("3.2 Payment terms are net thirty (30) days from the date of invoice. Late payments shall accrue interest at the rate of one and one-half percent (1.5%) per month or the maximum rate permitted by applicable law, whichever is less.", 2, 0.5),
    ]

    for text, level, indent in clauses:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(3)
        # Deliberately inconsistent indentation (simulating manual formatting)
        para.paragraph_format.left_indent = Inches(indent)
        # Set font for all runs
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    # Closing paragraph
    closing = doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.'
    )
    closing.paragraph_format.space_before = Pt(18)
    for run in closing.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
