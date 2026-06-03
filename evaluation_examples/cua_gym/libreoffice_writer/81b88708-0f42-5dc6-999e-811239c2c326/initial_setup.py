"""
Initial Setup: Wedding RSVP Card - pre-task state
Task ID: writer_creative_023
Domain: libreoffice_writer

Creates a basic RSVP draft file at ~/Desktop/wedding_rsvp.docx:
- Letter size, portrait orientation
- All text in 12pt Times New Roman, left-aligned
- No checkbox characters, no special formatting, no underline fills
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_023'
OUTPUT = f'{WORKDIR}/wedding_rsvp.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page to Letter size, portrait
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Lines for the RSVP card — all plain, 12pt Times New Roman, left-aligned
    lines = [
        'RSVP',
        'Kindly respond by May 15, 2026',
        'Name:',
        'Number of Guests:',
        'Joyfully Accepts',
        'Respectfully Declines',
    ]

    for line_text in lines:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(line_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = False
        run.italic = False

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
