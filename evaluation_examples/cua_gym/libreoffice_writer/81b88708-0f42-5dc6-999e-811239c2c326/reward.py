"""
Reward Script: Wedding RSVP Card Formatting
Task ID: writer_creative_023
Domain: libreoffice_writer
Scoring:
  - Component 1: Page setup A5 landscape (0.25 pts)
  - Component 2: RSVP line: 36pt, bold, centered (0.25 pts)
  - Component 3: 'Kindly respond...' line: 14pt, italic, centered (0.20 pts)
  - Component 4: Checkbox characters before Joyfully/Respectfully lines (0.15 pts)
  - Component 5: All paragraphs use Liberation Serif font family (0.15 pts)
"""

import os

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_023'

# A5 dimensions in mm (landscape: width=210, height=148)
A5_LANDSCAPE_WIDTH_MM = 210.0
A5_LANDSCAPE_HEIGHT_MM = 148.0
# Tolerance for page dimension checks (±3mm)
DIM_TOL_MM = 3.0


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ----------------------------------------------------------------
    # Component 1: Page setup — A5 landscape (0.25 points)
    # Initial env: Letter portrait (215.9x279.4mm)
    # Golden env: A5 landscape (210x148mm)
    # ----------------------------------------------------------------
    try:
        section = doc.sections[0]
        width_mm = section.page_width.mm
        height_mm = section.page_height.mm
        is_landscape = (section.orientation == WD_ORIENT.LANDSCAPE)

        width_ok = abs(width_mm - A5_LANDSCAPE_WIDTH_MM) <= DIM_TOL_MM
        height_ok = abs(height_mm - A5_LANDSCAPE_HEIGHT_MM) <= DIM_TOL_MM

        if is_landscape and width_ok and height_ok:
            print(f"PASS: Component 1 — A5 landscape page ({width_mm:.1f}x{height_mm:.1f}mm, orientation=LANDSCAPE) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected A5 landscape 210x148mm, found {width_mm:.1f}x{height_mm:.1f}mm, orientation={section.orientation}")
    except Exception as e:
        print(f"ERROR: Component 1 (page setup) — {e}")

    # ----------------------------------------------------------------
    # Component 2: 'RSVP' paragraph — 36pt, bold, centered (0.25 points)
    # Initial env: 12pt, not bold, left-aligned
    # ----------------------------------------------------------------
    try:
        rsvp_para = None
        for para in doc.paragraphs:
            if para.text.strip().upper() == 'RSVP':
                rsvp_para = para
                break

        if rsvp_para is None:
            print("FAIL: Component 2 — 'RSVP' paragraph not found")
        else:
            alignment_ok = (rsvp_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

            # Check runs for font size and bold using any() to avoid direct True assignment
            text_runs = [r for r in rsvp_para.runs if r.text.strip()]
            size_ok = any(r.font.size and abs(r.font.size.pt - 36.0) < 0.5 for r in text_runs)
            bold_ok = any(r.font.bold is True for r in text_runs)

            if alignment_ok and size_ok and bold_ok:
                print(f"PASS: Component 2 — 'RSVP' is 36pt, bold, centered (0.25 pts)")
                total_score += 0.25
            else:
                details = []
                if not alignment_ok:
                    details.append(f"alignment={rsvp_para.paragraph_format.alignment} (expected CENTER)")
                if not size_ok:
                    sizes = [run.font.size.pt if run.font.size else None for run in rsvp_para.runs if run.text.strip()]
                    details.append(f"size={sizes} (expected 36pt)")
                if not bold_ok:
                    details.append("not bold")
                print(f"FAIL: Component 2 — 'RSVP' formatting issues: {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 2 (RSVP formatting) — {e}")

    # ----------------------------------------------------------------
    # Component 3: 'Kindly respond...' — 14pt, italic, centered (0.20 points)
    # Initial env: 12pt, not italic, left-aligned
    # ----------------------------------------------------------------
    try:
        respond_para = None
        for para in doc.paragraphs:
            if 'Kindly respond' in para.text:
                respond_para = para
                break

        if respond_para is None:
            print("FAIL: Component 3 — 'Kindly respond by May 15, 2026' paragraph not found")
        else:
            alignment_ok = (respond_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

            # Check runs using any() to avoid direct True assignment
            text_runs = [r for r in respond_para.runs if r.text.strip()]
            size_ok = any(r.font.size and abs(r.font.size.pt - 14.0) < 0.5 for r in text_runs)
            italic_ok = any(r.font.italic is True for r in text_runs)

            if alignment_ok and size_ok and italic_ok:
                print(f"PASS: Component 3 — 'Kindly respond...' is 14pt, italic, centered (0.20 pts)")
                total_score += 0.20
            else:
                details = []
                if not alignment_ok:
                    details.append(f"alignment={respond_para.paragraph_format.alignment} (expected CENTER)")
                if not size_ok:
                    sizes = [run.font.size.pt if run.font.size else None for run in respond_para.runs if run.text.strip()]
                    details.append(f"size={sizes} (expected 14pt)")
                if not italic_ok:
                    details.append("not italic")
                print(f"FAIL: Component 3 — 'Kindly respond...' formatting issues: {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 3 (subtitle formatting) — {e}")

    # ----------------------------------------------------------------
    # Component 4: Checkbox characters (☐) before Joyfully/Respectfully (0.15 points)
    # Initial env: No checkbox characters, plain text
    # Golden env: '☐ Joyfully Accepts' and '☐ Respectfully Declines'
    # ----------------------------------------------------------------
    try:
        CHECKBOX = '\u2610'
        accepts_para = None
        declines_para = None
        for para in doc.paragraphs:
            text = para.text
            if 'Joyfully Accepts' in text:
                accepts_para = para
            if 'Respectfully Declines' in text:
                declines_para = para

        accepts_has_checkbox = accepts_para is not None and CHECKBOX in accepts_para.text
        declines_has_checkbox = declines_para is not None and CHECKBOX in declines_para.text

        if accepts_has_checkbox and declines_has_checkbox:
            print(f"PASS: Component 4 — Checkbox characters present before both RSVP options (0.15 pts)")
            total_score += 0.15
        else:
            details = []
            if not accepts_has_checkbox:
                val = repr(accepts_para.text) if accepts_para else "paragraph not found"
                details.append(f"'Joyfully Accepts' missing checkbox: {val}")
            if not declines_has_checkbox:
                val = repr(declines_para.text) if declines_para else "paragraph not found"
                details.append(f"'Respectfully Declines' missing checkbox: {val}")
            print(f"FAIL: Component 4 — {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 4 (checkbox characters) — {e}")

    # ----------------------------------------------------------------
    # Component 5: All text paragraphs use Liberation Serif font (0.15 points)
    # Initial env: Times New Roman throughout
    # Golden env: Liberation Serif throughout
    # ----------------------------------------------------------------
    try:
        EXPECTED_FONT = 'Liberation Serif'
        wrong_fonts = []
        checked_runs = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    checked_runs += 1
                    font_name = run.font.name
                    if font_name and font_name != EXPECTED_FONT:
                        wrong_fonts.append((para.text[:30], font_name))

        if checked_runs == 0:
            print("FAIL: Component 5 — No runs with text found to check font")
        elif not wrong_fonts:
            print(f"PASS: Component 5 — All {checked_runs} text runs use '{EXPECTED_FONT}' (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Some runs use wrong font: {wrong_fonts[:5]} (expected '{EXPECTED_FONT}')")
    except Exception as e:
        print(f"ERROR: Component 5 (font family) — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/Desktop/wedding_rsvp.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
