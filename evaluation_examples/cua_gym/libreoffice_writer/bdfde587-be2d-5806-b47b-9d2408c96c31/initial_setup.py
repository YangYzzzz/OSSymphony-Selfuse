"""
Initial Setup: Create a court filing document (motion for summary judgment) with no footer.
Task ID: writer_legal_023
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_023'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_centered_heading(doc, text, size=14, bold=True, space_after=6):
    """Add a centered, bold heading paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return para


def add_body_para(doc, text, first_indent=True, space_after=6, bold=False):
    """Add a justified body paragraph with optional first-line indent."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return para


def add_numbered_para(doc, number, text, space_after=6):
    """Add a numbered paragraph for legal arguments."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    run = para.add_run(f"{number}. ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = para.add_run(text)
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    return para


def create_initial():
    doc = Document()

    # Page setup: letter size with 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -- Court caption --
    add_centered_heading(doc, 'IN THE SUPERIOR COURT OF THE STATE OF CALIFORNIA', size=12)
    add_centered_heading(doc, 'FOR THE COUNTY OF LOS ANGELES', size=12)

    # Parties block
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run('WESTFIELD COMMERCIAL PROPERTIES, LLC,')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.left_indent = Inches(2)
    run = para.add_run('Plaintiff,')
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = para.add_run('v.')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run('PACIFIC RIM CONSTRUCTION, INC., and')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run('DAVID NAKAMURA, individually,')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.left_indent = Inches(2)
    run = para.add_run('Defendants.')
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Case number line (in body, NOT in footer)
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run('Case No. 2024-CV-03891')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # -- Title of the motion --
    add_centered_heading(doc, 'MOTION FOR SUMMARY JUDGMENT', size=14)
    add_centered_heading(doc, 'OR, IN THE ALTERNATIVE, SUMMARY ADJUDICATION', size=12, bold=True, space_after=12)

    # -- Introduction --
    add_body_para(doc,
        'Plaintiff Westfield Commercial Properties, LLC ("Plaintiff" or "Westfield"), '
        'by and through its attorneys of record, hereby moves this Court pursuant to '
        'California Code of Civil Procedure Section 437c for summary judgment, or in '
        'the alternative, summary adjudication of issues, against Defendants Pacific Rim '
        'Construction, Inc. ("Pacific Rim") and David Nakamura ("Nakamura") (collectively, '
        '"Defendants"). This motion is made on the grounds that there are no triable issues '
        'of material fact and Plaintiff is entitled to judgment as a matter of law.',
        space_after=12)

    # -- Statement of undisputed facts --
    add_centered_heading(doc, 'STATEMENT OF UNDISPUTED MATERIAL FACTS', size=12, space_after=10)

    add_numbered_para(doc, 1,
        'On or about March 15, 2022, Westfield entered into a written construction '
        'agreement (the "Agreement") with Pacific Rim for the renovation of the '
        'commercial property located at 4500 Wilshire Boulevard, Suite 200, '
        'Los Angeles, California 90010 (the "Property"). (Declaration of Robert '
        'Thornton, \u00B6 3, Exhibit A.)')

    add_numbered_para(doc, 2,
        'The Agreement specified a total contract price of $2,875,000.00 for all '
        'renovation work, with a completion deadline of December 31, 2022. '
        '(Thornton Decl., \u00B6 4, Exhibit A at pp. 3-4.)')

    add_numbered_para(doc, 3,
        'Defendant Nakamura personally guaranteed Pacific Rim\'s performance '
        'under the Agreement by executing a written personal guaranty dated '
        'March 15, 2022 (the "Guaranty"). (Thornton Decl., \u00B6 5, Exhibit B.)')

    add_numbered_para(doc, 4,
        'Westfield made progress payments to Pacific Rim totaling $1,950,000.00 '
        'between April 2022 and October 2022, in accordance with the payment '
        'schedule set forth in the Agreement. (Thornton Decl., \u00B6\u00B6 6-8, '
        'Exhibit C.)')

    add_numbered_para(doc, 5,
        'Pacific Rim ceased all work on the Property on or about November 10, 2022, '
        'when the renovation was approximately 55% complete. Pacific Rim did not '
        'provide notice of work stoppage as required by Section 12.3 of the Agreement. '
        '(Thornton Decl., \u00B6 9; Declaration of Maria Santos, \u00B6\u00B6 4-6.)')

    add_numbered_para(doc, 6,
        'An independent construction expert, James Richardson, P.E., inspected the '
        'Property on January 18, 2023, and determined that the reasonable value of '
        'work completed by Pacific Rim was $1,150,000.00. (Declaration of James '
        'Richardson, \u00B6\u00B6 7-12, Exhibit D.)')

    add_numbered_para(doc, 7,
        'Westfield incurred additional costs of $2,100,000.00 to retain a replacement '
        'contractor, Meridian Building Group, Inc., to complete the remaining renovation '
        'work. (Thornton Decl., \u00B6\u00B6 14-16, Exhibit E.)')

    # -- Legal argument --
    add_centered_heading(doc, 'MEMORANDUM OF POINTS AND AUTHORITIES', size=12, space_after=10)

    para = add_body_para(doc, 'I. STANDARD OF REVIEW', first_indent=False, bold=True, space_after=8)

    add_body_para(doc,
        'A party is entitled to summary judgment where "all the papers submitted show '
        'that there is no triable issue as to any material fact and that the moving party '
        'is entitled to a judgment as a matter of law." Cal. Code Civ. Proc. \u00A7 437c(c). '
        'The moving party bears the initial burden of production to make a prima facie '
        'showing that no triable issue of material fact exists. Aguilar v. Atlantic '
        'Richfield Co. (2001) 25 Cal.4th 826, 850.',
        space_after=10)

    para = add_body_para(doc, 'II. PLAINTIFF IS ENTITLED TO SUMMARY JUDGMENT ON ITS BREACH OF CONTRACT CLAIM',
                         first_indent=False, bold=True, space_after=8)

    add_body_para(doc,
        'To establish a claim for breach of contract under California law, a plaintiff '
        'must prove: (1) the existence of a contract; (2) the plaintiff\'s performance or '
        'excuse for nonperformance; (3) the defendant\'s breach; and (4) resulting damages. '
        'Oasis West Realty, LLC v. Goldman (2011) 51 Cal.4th 811, 821.',
        space_after=8)

    add_body_para(doc,
        'Here, each element is established by undisputed evidence. The Agreement '
        'constitutes a valid and enforceable contract between Westfield and Pacific Rim. '
        '(UMF \u00B6 1.) Westfield fully performed its obligations by making timely progress '
        'payments totaling $1,950,000.00. (UMF \u00B6 4.) Pacific Rim breached the Agreement '
        'by abandoning the project before completion without proper notice. (UMF \u00B6 5.) '
        'Westfield suffered damages in the amount of $2,900,000.00, representing the '
        'difference between amounts paid ($1,950,000.00) and the value of work received '
        '($1,150,000.00), plus the additional cost to complete ($2,100,000.00). '
        '(UMF \u00B6\u00B6 4, 6, 7.)',
        space_after=10)

    para = add_body_para(doc, 'III. NAKAMURA IS LIABLE UNDER THE PERSONAL GUARANTY',
                         first_indent=False, bold=True, space_after=8)

    add_body_para(doc,
        'A guaranty is an agreement to answer for the debt or obligation of another. '
        'Cal. Civ. Code \u00A7 2787. Where a guarantor unconditionally guarantees the '
        'performance of the principal\'s obligations, the guarantor is liable for all '
        'damages flowing from the principal\'s breach. Gray1 CPB, LLC v. Kolokowsky '
        '(2011) 193 Cal.App.4th 803, 809.',
        space_after=8)

    add_body_para(doc,
        'Nakamura executed an unconditional personal guaranty of Pacific Rim\'s '
        'obligations under the Agreement. (UMF \u00B6 3.) The Guaranty provides that '
        '"Guarantor unconditionally guarantees to Owner the full and faithful performance '
        'by Contractor of each and every obligation of Contractor under the Agreement." '
        '(Thornton Decl., Exhibit B at \u00A7 1.) Because Pacific Rim breached the Agreement, '
        'Nakamura is jointly and severally liable for all resulting damages.',
        space_after=12)

    # -- Conclusion --
    add_centered_heading(doc, 'CONCLUSION', size=12, space_after=10)

    add_body_para(doc,
        'For the foregoing reasons, Plaintiff Westfield Commercial Properties, LLC '
        'respectfully requests that this Court grant summary judgment in Plaintiff\'s '
        'favor and against Defendants Pacific Rim Construction, Inc. and David Nakamura, '
        'jointly and severally, in the amount of $2,900,000.00, plus prejudgment interest, '
        'attorneys\' fees, and costs of suit.',
        space_after=18)

    # Signature block
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run('Dated: January 22, 2024')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # blank line

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = para.add_run('MORRISON & FITZGERALD LLP')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = para.add_run('By: ___________________________')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = para.add_run('Elizabeth A. Morrison, Esq.')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = para.add_run('Attorneys for Plaintiff')
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # NO FOOTER - this is what the task asks the agent to add

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
