"""
Reward Script: Create exhibit table in Master Services Agreement
Task ID: writer_legal_029
Domain: libreoffice_writer
Scoring:
  Component 1: Table exists with 5 rows x 3 cols (0.25)
  Component 2: Header row has correct column names (0.20)
  Component 3: Header row text is bold (0.15)
  Component 4: Data rows have correct exhibit letters A-D (0.20)
  Component 5: Data rows have correct exhibit titles (0.20)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_029'


def persist_app_state(domain: str):
    """Save any unsaved changes in LibreOffice Writer."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: file must have at least one table (task adds a table)
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document. Task requires adding an exhibit table.")
        print("REWARD: 0.0")
        return 0.0

    # Find the exhibit table - look for a table with 'Exhibit' in header
    exhibit_table = None
    for table in doc.tables:
        header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if 'exhibit' in header_texts:
            exhibit_table = table
            break

    if exhibit_table is None:
        # Fallback: use the last table added
        exhibit_table = doc.tables[-1]

    # Component 1: Table has correct dimensions - 5 rows x 3 cols (0.25 points)
    try:
        num_rows = len(exhibit_table.rows)
        num_cols = len(exhibit_table.columns)
        if num_rows == 5 and num_cols == 3:
            print(f"PASS: Component 1 - Table has correct dimensions 5x3 (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - Expected 5 rows x 3 cols, found {num_rows} rows x {num_cols} cols")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Header row has correct column names (0.20 points)
    try:
        header_cells = [cell.text.strip() for cell in exhibit_table.rows[0].cells]
        expected_headers = ['Exhibit', 'Title', 'Page']
        # Case-insensitive comparison
        headers_match = [h.lower() for h in header_cells] == [h.lower() for h in expected_headers]
        if headers_match:
            print(f"PASS: Component 2 - Header row has correct column names {header_cells} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 - Expected headers {expected_headers}, found {header_cells}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Header row text is bold (0.15 points)
    try:
        header_bold_count = 0
        for ci in range(min(3, len(exhibit_table.columns))):
            cell = exhibit_table.rows[0].cells[ci]
            cell_has_bold = False
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip() and run.font.bold:
                        cell_has_bold = True
                        break
                if cell_has_bold:
                    break
            if cell_has_bold:
                header_bold_count += 1

        if header_bold_count == 3:
            print(f"PASS: Component 3 - All 3 header cells have bold text (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 - Only {header_bold_count}/3 header cells have bold text")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Data rows have correct exhibit letters A, B, C, D (0.20 points)
    try:
        expected_letters = ['A', 'B', 'C', 'D']
        actual_letters = []
        for ri in range(1, min(5, len(exhibit_table.rows))):
            cell_text = exhibit_table.rows[ri].cells[0].text.strip().upper()
            actual_letters.append(cell_text)

        matches = sum(1 for exp, act in zip(expected_letters, actual_letters) if exp == act)
        if matches == 4:
            print(f"PASS: Component 4 - All 4 exhibit letters correct {actual_letters} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 - Expected {expected_letters}, found {actual_letters} ({matches}/4 match)")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Data rows have correct exhibit titles (0.20 points)
    try:
        expected_titles = [
            'Scope of Services',
            'Fee Schedule',
            'Insurance Requirements',
            'Service Level Agreement'
        ]
        actual_titles = []
        for ri in range(1, min(5, len(exhibit_table.rows))):
            cell_text = exhibit_table.rows[ri].cells[1].text.strip()
            actual_titles.append(cell_text)

        title_matches = 0
        for exp, act in zip(expected_titles, actual_titles):
            if exp.lower() == act.lower():
                title_matches += 1

        if title_matches == 4:
            print(f"PASS: Component 5 - All 4 exhibit titles correct (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 5 - {title_matches}/4 titles match. Expected {expected_titles}, found {actual_titles}")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
