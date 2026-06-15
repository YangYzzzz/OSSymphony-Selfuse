"""
Initial Setup: Create a master services agreement document that references exhibits but has no exhibit table.
Task ID: writer_legal_029
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('MASTER SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(6)
    run = preamble.add_run(
        'This Master Services Agreement ("Agreement") is entered into as of March 15, 2025 '
        '("Effective Date"), by and between Northwind Technologies, Inc., a Delaware corporation '
        'with principal offices at 4200 Lakewood Boulevard, Suite 300, Denver, CO 80246 ("Client"), '
        'and Pinnacle Consulting Group, LLC, a California limited liability company with principal '
        'offices at 1580 Innovation Drive, San Jose, CA 95134 ("Service Provider").'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Recitals ---
    recitals_heading = doc.add_heading('RECITALS', level=1)

    recitals_items = [
        'WHEREAS, Client desires to engage Service Provider to perform certain professional '
        'services as described in Exhibit A (Scope of Services) attached hereto;',
        'WHEREAS, Service Provider represents that it has the expertise, resources, and '
        'qualifications necessary to perform such services in accordance with industry standards;',
        'WHEREAS, the parties wish to establish the terms and conditions under which Service '
        'Provider will render services to Client, including the fee structure set forth in '
        'Exhibit B (Fee Schedule);',
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements contained '
        'herein, and for other good and valuable consideration, the receipt and sufficiency '
        'of which are hereby acknowledged, the parties agree as follows:'
    ]

    for item in recitals_items:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(6)

    # --- Article 1: Definitions ---
    doc.add_heading('ARTICLE 1: DEFINITIONS', level=1)

    definitions = [
        ('"Confidential Information"', ' means any non-public information disclosed by either '
         'party to the other, whether orally, in writing, or by inspection, including but not '
         'limited to trade secrets, business plans, financial data, customer lists, and '
         'technical specifications.'),
        ('"Deliverables"', ' means the tangible and intangible work product to be delivered '
         'by Service Provider as specified in Exhibit A (Scope of Services).'),
        ('"Service Level Agreement" or "SLA"', ' means the performance standards and metrics '
         'set forth in Exhibit D (Service Level Agreement) attached hereto.'),
        ('"Term"', ' means the period commencing on the Effective Date and continuing for '
         'twenty-four (24) months unless earlier terminated in accordance with Article 7.'),
    ]

    for bold_part, normal_part in definitions:
        p = doc.add_paragraph()
        r1 = p.add_run(bold_part)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(normal_part)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(6)

    # --- Article 2: Scope of Services ---
    doc.add_heading('ARTICLE 2: SCOPE OF SERVICES', level=1)

    scope_paras = [
        'Service Provider shall perform the services described in Exhibit A (Scope of Services) '
        'in a professional and workmanlike manner consistent with generally accepted industry '
        'standards and practices.',
        'Any modifications to the scope of services must be agreed upon in writing by both '
        'parties through a formal change order process as outlined in Section 2.3 below.',
        'Service Provider shall maintain adequate staffing levels to meet the obligations set '
        'forth in this Agreement and shall comply with the insurance requirements detailed in '
        'Exhibit C (Insurance Requirements).',
    ]

    for text in scope_paras:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(6)

    # --- Article 3: Compensation ---
    doc.add_heading('ARTICLE 3: COMPENSATION', level=1)

    comp_paras = [
        'Client shall compensate Service Provider in accordance with the fee structure set forth '
        'in Exhibit B (Fee Schedule). All invoices shall be submitted monthly and are due within '
        'thirty (30) days of receipt.',
        'Service Provider shall maintain detailed time records and expense reports. Travel expenses '
        'exceeding $500.00 per trip must be pre-approved in writing by Client.',
        'Late payments shall accrue interest at a rate of 1.5% per month or the maximum rate '
        'permitted by applicable law, whichever is less.',
    ]

    for text in comp_paras:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(6)

    # --- Article 4: Confidentiality ---
    doc.add_heading('ARTICLE 4: CONFIDENTIALITY', level=1)

    conf_paras = [
        'Each party agrees to hold the other party\'s Confidential Information in strict '
        'confidence and not to disclose such information to any third party without the prior '
        'written consent of the disclosing party.',
        'The obligations of confidentiality shall survive the termination or expiration of '
        'this Agreement for a period of five (5) years.',
    ]

    for text in conf_paras:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(6)

    # --- Article 5: Performance Standards ---
    doc.add_heading('ARTICLE 5: PERFORMANCE STANDARDS', level=1)

    perf_paras = [
        'Service Provider shall adhere to the performance metrics and service levels specified '
        'in Exhibit D (Service Level Agreement). Failure to meet the agreed-upon service levels '
        'may result in service credits as described therein.',
        'Client shall have the right to conduct periodic performance reviews to assess Service '
        'Provider\'s compliance with the terms of this Agreement and the standards set forth '
        'in the applicable Exhibits.',
    ]

    for text in perf_paras:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(6)

    # --- Signature block ---
    doc.add_paragraph()  # spacer
    sig_heading = doc.add_heading('SIGNATURES', level=1)

    sig_text = (
        'IN WITNESS WHEREOF, the parties have executed this Master Services Agreement '
        'as of the Effective Date first written above.'
    )
    p = doc.add_paragraph()
    r = p.add_run(sig_text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(24)

    # Client signature
    for line in ['NORTHWIND TECHNOLOGIES, INC.', '', 'By: _______________________________',
                 'Name: Jennifer A. Morrison', 'Title: Chief Operating Officer',
                 'Date: _______________________________']:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        if line.startswith('NORTHWIND'):
            r.bold = True
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacer

    for line in ['PINNACLE CONSULTING GROUP, LLC', '', 'By: _______________________________',
                 'Name: David R. Takahashi', 'Title: Managing Partner',
                 'Date: _______________________________']:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        if line.startswith('PINNACLE'):
            r.bold = True
        p.paragraph_format.space_after = Pt(2)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
