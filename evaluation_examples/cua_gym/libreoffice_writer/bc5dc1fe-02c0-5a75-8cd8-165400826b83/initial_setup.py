"""
Initial Setup: Open LibreOffice Writer with a sample document.
AutoCorrect settings are at default (no 'iPhone' exception).
Task ID: writer_frd_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_064'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    doc.add_heading("Product Launch Brief — Q3 2025", level=1)

    # Intro paragraph
    para = doc.add_paragraph(
        "The upcoming product launch for the iPhone 16 Pro Max is scheduled for "
        "September 2025. The marketing team needs to finalize all campaign materials "
        "by August 15th to allow for review and approval cycles."
    )

    # Section: Key Milestones
    doc.add_heading("Key Milestones", level=2)
    milestones = [
        "July 1 — Creative brief finalized and approved by stakeholders",
        "July 15 — Photography and video assets delivered from studio",
        "August 1 — Digital ad campaign configured in Google Ads and Meta",
        "August 15 — All materials submitted for legal and compliance review",
        "September 5 — iPhone pre-order page goes live on website",
        "September 12 — Official launch event and press release",
    ]
    for m in milestones:
        doc.add_paragraph(m, style="List Bullet")

    # Section: Budget Overview
    doc.add_heading("Budget Overview", level=2)
    budget_para = doc.add_paragraph(
        "The total allocated budget for the iPhone launch campaign is $4.2 million. "
        "This includes $1.8M for digital advertising, $900K for influencer partnerships, "
        "$750K for in-store displays, and $750K for PR and events."
    )

    # Section: Team Responsibilities
    doc.add_heading("Team Responsibilities", level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"

    headers = ["Team Member", "Role", "Deliverable"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    team_data = [
        ["Sarah Chen", "Campaign Lead", "Overall coordination and timeline management"],
        ["Marcus Rivera", "Creative Director", "Ad creatives, video scripts, brand guidelines"],
        ["Priya Sharma", "Digital Marketing", "Google Ads, social media, email campaigns"],
        ["James O'Brien", "PR Manager", "Press releases, media outreach, launch event"],
        ["Aiko Tanaka", "Analytics Lead", "KPI tracking, A/B testing, post-launch reporting"],
    ]
    for r, row_data in enumerate(team_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Notes section
    doc.add_heading("Notes", level=2)
    doc.add_paragraph(
        "Remember that when typing iPhone at the beginning of a sentence in this document, "
        "LibreOffice may auto-capitalize it to IPhone. This should be corrected by adding "
        "an AutoCorrect exception."
    )
    doc.add_paragraph(
        "The iPhone branding must always use a lowercase 'i' followed by an uppercase 'P'. "
        "This is a trademark requirement from Apple Inc."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
