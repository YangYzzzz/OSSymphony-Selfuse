"""
Reward Script: Children's Fractions Worksheet in LibreOffice Writer
Task ID: writer_wf_052
Domain: libreoffice_writer
Scoring:
  C1: Title 'Fun with Fractions - Grade 4' in 18pt bold (0.15)
  C2: Name and Date fields present (0.10)
  C3: 5 fraction addition problems with answer blanks (0.20)
  C4: 5 fraction comparison problems using < > = (0.20)
  C5: 2 word problems with answers (0.15)
  C6: Bonus challenge section (0.10)
  C7: Footer with Page ___ of ___ field codes (0.10)
"""

import os
import re
from docx import Document
from docx.shared import Pt
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_052'

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_texts = [p.text.strip() for p in doc.paragraphs]
    all_texts_lower = [t.lower() for t in all_texts]

    # Component 1: Title 'Fun with Fractions - Grade 4' in 18pt bold (0.15 points)
    try:
        title_found = False
        for p in doc.paragraphs:
            text = p.text.strip()
            if 'fun with fractions' in text.lower() and 'grade 4' in text.lower():
                # Check for bold and ~18pt in at least one run
                for r in p.runs:
                    if r.text.strip():
                        is_bold = r.font.bold is True
                        size_ok = (r.font.size is not None and
                                   abs(r.font.size.pt - 18.0) < 1.0)
                        if is_bold and size_ok:
                            title_found = True
                            break
                if title_found:
                    break
        if title_found:
            print(f"PASS: Component 1 — Title found with 18pt bold (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title 'Fun with Fractions - Grade 4' not found in 18pt bold")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Name and Date fields present (0.10 points)
    try:
        has_name = False
        has_date = False
        for t in all_texts_lower:
            if 'name' in t and ('_' in t or 'blank' in t or ':' in t):
                has_name = True
            if 'date' in t and ('_' in t or 'blank' in t or ':' in t):
                has_date = True
        if has_name and has_date:
            print(f"PASS: Component 2 — Name and Date fields found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Name={has_name}, Date={has_date}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 5 fraction addition problems with answer blanks (0.20 points)
    # Look for lines with fraction patterns like a/b + c/d
    try:
        addition_count = 0
        fraction_add_pattern = re.compile(r'\d+\s*/\s*\d+\s*\+\s*\d+\s*/\s*\d+')
        for t in all_texts:
            if fraction_add_pattern.search(t) and '_' in t:
                addition_count += 1
        if addition_count >= 5:
            print(f"PASS: Component 3 — {addition_count} fraction addition problems found (0.20 pts)")
            total_score += 0.20
        elif addition_count >= 3:
            partial = 0.10
            print(f"PARTIAL: Component 3 — {addition_count}/5 addition problems found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {addition_count} fraction addition problems found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 5 fraction comparison problems using < > = (0.20 points)
    # Look for lines with two fractions and a blank between them (for < > =)
    try:
        comparison_count = 0
        # Pattern: fraction blank fraction (the blank is for < > =)
        frac_pattern = re.compile(r'\d+\s*/\s*\d+')
        for t in all_texts:
            fracs = frac_pattern.findall(t)
            # Must have exactly 2 fractions and a blank/underscore between them, no + sign
            if len(fracs) >= 2 and '+' not in t and '_' in t:
                comparison_count += 1
        if comparison_count >= 5:
            print(f"PASS: Component 4 — {comparison_count} fraction comparison problems found (0.20 pts)")
            total_score += 0.20
        elif comparison_count >= 3:
            partial = 0.10
            print(f"PARTIAL: Component 4 — {comparison_count}/5 comparison problems found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {comparison_count} fraction comparison problems found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Word problem section with 2 problems (0.15 points)
    try:
        # Look for section header mentioning word problems
        has_word_section = any('word problem' in t for t in all_texts_lower)
        # Count word problems: longer text paragraphs containing fractions in prose context
        word_problem_count = 0
        in_word_section = False
        for t in all_texts:
            tl = t.lower()
            if 'word problem' in tl:
                in_word_section = True
                continue
            if in_word_section and 'bonus' in tl:
                break
            if in_word_section and len(t) > 40 and '/' in t:
                word_problem_count += 1

        if has_word_section and word_problem_count >= 2:
            print(f"PASS: Component 5 — Word problems section with {word_problem_count} problems (0.15 pts)")
            total_score += 0.15
        elif has_word_section and word_problem_count >= 1:
            partial = 0.08
            print(f"PARTIAL: Component 5 — Word section found but only {word_problem_count} problem(s) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Word section={has_word_section}, problems={word_problem_count}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Bonus challenge section (0.10 points)
    try:
        has_bonus = any('bonus' in t for t in all_texts_lower)
        # Bonus should have some content after it
        bonus_has_content = False
        found_bonus = False
        for t in all_texts:
            if found_bonus and len(t.strip()) > 10:
                bonus_has_content = True
                break
            if 'bonus' in t.lower():
                found_bonus = True

        if has_bonus and bonus_has_content:
            print(f"PASS: Component 6 — Bonus challenge section found with content (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — Bonus={has_bonus}, content={bonus_has_content}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Footer with 'Page ___ of ___' using PAGE/NUMPAGES field codes (0.10 points)
    try:
        footer_ok = False
        for section in doc.sections:
            if section.footer and section.footer.paragraphs:
                footer_xml = etree.tostring(section.footer._element).decode()
                has_page_text = 'Page' in footer_xml or 'page' in footer_xml
                has_page_field = 'PAGE' in footer_xml
                has_numpages_field = 'NUMPAGES' in footer_xml
                has_of = ' of ' in footer_xml

                if has_page_text and has_page_field and has_numpages_field and has_of:
                    footer_ok = True
                    break

        if footer_ok:
            print(f"PASS: Component 7 — Footer with 'Page ___ of ___' field codes found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 — Footer missing or lacks PAGE/NUMPAGES fields")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
