"""
Reward Script: HR Department Meeting Minutes
Task ID: writer_hr_035
Domain: libreoffice_writer
Scoring:
  C1 (0.15): Header "HR Department Meeting Minutes" present
  C2 (0.10): Date line with "March 15, 2026"
  C3 (0.15): Attendees list with bullet items (at least 3)
  C4 (0.25): 5 numbered agenda items with discussion notes
  C5 (0.20): Action items table with 3 columns (Action, Owner, Due Date) and >= 3 data rows
  C6 (0.15): "Next Meeting" line at the bottom
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_035'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = "\n".join(p.text for p in doc.paragraphs)

    # Component 1: Header "HR Department Meeting Minutes" (0.15 points)
    # The initial doc is blank, so any heading with this text is task-introduced.
    try:
        header_found = False
        for p in doc.paragraphs:
            if "hr department meeting minutes" in p.text.lower().strip():
                header_found = True
                break
        if header_found:
            print(f"PASS: Component 1 - Header 'HR Department Meeting Minutes' found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 - Header 'HR Department Meeting Minutes' not found")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Date line with "March 15, 2026" (0.10 points)
    try:
        date_found = False
        for p in doc.paragraphs:
            if "march 15, 2026" in p.text.lower() or "march 15 2026" in p.text.lower():
                date_found = True
                break
        if date_found:
            print(f"PASS: Component 2 - Date 'March 15, 2026' found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 - Date 'March 15, 2026' not found in any paragraph")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Attendees list with at least 3 bullet items (0.15 points)
    # Initial doc has no bullet items, so any bullet items are task-introduced.
    try:
        bullet_count = 0
        found_attendees_label = False
        for p in doc.paragraphs:
            text_lower = p.text.lower().strip()
            if "attendee" in text_lower:
                found_attendees_label = True
            style_name = p.style.name.lower() if p.style else ""
            if "bullet" in style_name or "list" in style_name:
                bullet_count += 1

        if found_attendees_label and bullet_count >= 3:
            print(f"PASS: Component 3 - Attendees label found and {bullet_count} bullet items (0.15 pts)")
            total_score += 0.15
        elif bullet_count >= 3:
            # Some implementations may not use a separate "Attendees:" label but still list people
            # Give partial: check if bullets contain name-like content
            print(f"PASS: Component 3 - {bullet_count} bullet/list items found (attendees implied) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 - Attendees label={found_attendees_label}, bullet items={bullet_count} (need >= 3)")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: 5 numbered agenda items with discussion notes (0.25 points)
    # Check for paragraphs starting with "1.", "2.", etc. up to "5."
    # Each numbered item should have at least one follow-up paragraph (discussion note).
    try:
        agenda_items_found = 0
        agenda_with_notes = 0
        paragraphs = doc.paragraphs
        for i, p in enumerate(paragraphs):
            text = p.text.strip()
            # Match numbered items like "1. ...", "2. ...", etc.
            match = re.match(r'^(\d+)\.\s+\S', text)
            if match:
                num = int(match.group(1))
                if 1 <= num <= 5:
                    agenda_items_found += 1
                    # Check if next paragraph has substantial text (discussion note)
                    if i + 1 < len(paragraphs):
                        next_text = paragraphs[i + 1].text.strip()
                        if len(next_text) > 20:  # non-trivial discussion note
                            agenda_with_notes += 1

        if agenda_items_found >= 5 and agenda_with_notes >= 4:
            print(f"PASS: Component 4 - {agenda_items_found} agenda items, {agenda_with_notes} with discussion notes (0.25 pts)")
            total_score += 0.25
        elif agenda_items_found >= 3:
            partial = round(0.25 * (agenda_items_found / 5), 2)
            print(f"PARTIAL: Component 4 - {agenda_items_found}/5 agenda items, {agenda_with_notes} with notes ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - Found {agenda_items_found} agenda items, {agenda_with_notes} with notes (need 5)")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Action items table with 3 columns and >= 3 data rows (0.20 points)
    # Initial doc has 0 tables, so any table is task-introduced.
    try:
        table_found = False
        correct_headers = False
        sufficient_rows = False

        if len(doc.tables) >= 1:
            table_found = True
            # Check the first table (or find one with Action/Owner/Due Date headers)
            for table in doc.tables:
                if len(table.columns) >= 3:
                    header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    if "action" in header_cells and "owner" in header_cells and "due" in " ".join(header_cells):
                        correct_headers = True
                        data_rows = len(table.rows) - 1  # exclude header
                        if data_rows >= 3:
                            sufficient_rows = True
                        break

        if table_found and correct_headers and sufficient_rows:
            print(f"PASS: Component 5 - Action items table with correct headers and {data_rows} data rows (0.20 pts)")
            total_score += 0.20
        elif table_found and correct_headers:
            print(f"PARTIAL: Component 5 - Table has correct headers but only {data_rows} data rows (0.10 pts)")
            total_score += 0.10
        elif table_found:
            print(f"PARTIAL: Component 5 - Table exists but headers don't match (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 - No table found in document")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: "Next Meeting" line at the bottom (0.15 points)
    try:
        next_meeting_found = False
        # Check last 5 paragraphs for "next meeting" text
        last_paras = [p for p in doc.paragraphs if p.text.strip()]
        if last_paras:
            for p in last_paras[-5:]:
                if "next meeting" in p.text.lower():
                    next_meeting_found = True
                    break
        if next_meeting_found:
            print(f"PASS: Component 6 - 'Next Meeting' line found near bottom (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 6 - 'Next Meeting' line not found in last paragraphs")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
