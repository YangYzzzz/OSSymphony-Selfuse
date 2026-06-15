"""
Initial Setup: Quarterly Sales Data Analysis Report - initial state (all paragraphs left-aligned)
Task ID: writer_para_045
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1 - must remain unchanged (no explicit alignment set)
    heading = doc.add_heading('Quarterly Sales Data Analysis Report', level=1)
    # Leave heading alignment as default (inherit from style)

    # Paragraph 2: Table header - LEFT aligned (task requires centering this)
    para2 = doc.add_paragraph('Table 1: Regional Sales Performance Q4 2024')
    para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Paragraph 3: Data description - LEFT aligned (task requires justifying this)
    para3 = doc.add_paragraph(
        'The North American region led sales performance with $4.2 million in revenue, '
        'representing a 15% increase over the previous quarter. This growth was primarily '
        'driven by the successful launch of three new product lines in October.'
    )
    para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Paragraph 4: Table header - LEFT aligned (task requires centering this)
    para4 = doc.add_paragraph('Table 2: Product Category Breakdown')
    para4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Paragraph 5: Data description - LEFT aligned (task requires justifying this)
    para5 = doc.add_paragraph(
        'Enterprise solutions accounted for 62% of total revenue, while the consumer segment '
        'showed the fastest growth rate at 28% quarter-over-quarter. The professional services '
        'category remained stable at $800,000.'
    )
    para5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Paragraph 6: Table header - LEFT aligned (task requires centering this)
    para6 = doc.add_paragraph('Table 3: Customer Acquisition Metrics')
    para6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Paragraph 7: Data description - LEFT aligned (task requires justifying this)
    para7 = doc.add_paragraph(
        'New customer acquisition cost decreased to $340 per customer from $425 in the prior '
        'quarter. The improvement is attributed to the optimized digital marketing funnel '
        'implemented in September.'
    )
    para7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open with LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
