"""
FINAL REWARD SCRIPT - SUCCESS
Task: I just noticed I skipped an item and need it near the top of my list. In LibreOffice Writer, how can I slip one extra row right above row 2 in “Table 1” so everything else shifts down neatly?
Generated: 2025-09-10 16:36:40
Status: success
Model: azure-o3
Total Steps: 2
"""

from docx import Document
import os, re

def verify_task(file_path: str) -> float:
    """Verify that the user inserted a new row directly above the old second
    row in Table 1 so every subsequent row shifted down.
    Progressive scoring awards points for:
      1. Detecting Table 1 with at least 2 rows (0.4)
      2. Verifying that the original first item ("Item 1") is still on row 0
         and that the original second item ("Item 2") is now on row ≥ 2,
         proving a row was inserted above it (0.15 + 0.15 = 0.30)
      3. Detecting an empty (new) row exactly at index 1, right above the old
         second row (0.30)
    The script returns a float between 0 and 1 and prints "REWARD: X.X".
    """

    score = 0.0
    max_score = 1.0
    print(f"Verifying file: {file_path}")

    # ---------- 1. File existence and loading ----------
    if not os.path.exists(file_path):
        print("✗ File not found")
        return 0.0  # No points for missing file

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Failed to load document: {e}")
        return 0.0  # Loading must succeed to continue

    # ---------- 2. Locate Table 1 ----------
    if not doc.tables:
        print("✗ No tables found in document")
        return 0.0

    table = doc.tables[0]
    rows = len(table.rows)
    cols = len(table.columns)
    print(f"✓ Found Table 1 with {rows} rows and {cols} columns")

    # Require at least two rows to assess insertion
    if rows >= 2:
        score += 0.4  # Table with sufficient rows detected

    # ---------- 3. Gather text from each row ----------
    row_texts = [
        " ".join(cell.text.strip() for cell in row.cells).strip()
        for row in table.rows
    ]

    for idx, txt in enumerate(row_texts):
        print(f"Row {idx}: '{txt}'")

    # ---------- 4. Verify positions of original items ----------
    idx_item1 = next((i for i, txt in enumerate(row_texts) if re.fullmatch(r"Item 1", txt)), None)
    idx_item2 = next((i for i, txt in enumerate(row_texts) if re.fullmatch(r"Item 2", txt)), None)

    if idx_item1 is not None and idx_item2 is not None:
        print(f"Found 'Item 1' at row {idx_item1}, 'Item 2' at row {idx_item2}")

        # Original first item should remain on top
        if idx_item1 == 0:
            score += 0.15

        # Original second item should now be at row ≥ 2 (shifted down)
        if idx_item2 >= 2:
            print("✓ Detected that an extra row exists above original row 2")
            score += 0.15
        else:
            print("✗ 'Item 2' was not shifted down; row insertion likely missing")
    else:
        print("✗ Could not locate both 'Item 1' and 'Item 2' for verification")

    # ---------- 5. Check for a blank row precisely at index 1 ----------
    blank_row_idx = None
    for i in range(min(3, rows)):
        if not row_texts[i]:  # Completely empty text → inserted blank row
            blank_row_idx = i
            break

    if blank_row_idx is not None and blank_row_idx == 1:
        print("✓ Blank row detected at index 1 (right above original row 2)")
        score += 0.30
    else:
        print("✗ No blank row found at index 1; cannot confirm correct insertion")

    # ---------- 6. Final score ----------
    final_score = min(score, max_score)
    print(f"Total score: {final_score}")
    print(f"REWARD: {final_score}")
    return final_score

# --------- Execute verification on provided file ---------
if __name__ == "__main__":
    FILE_PATH = "/home/user/i_just_noticed_i_skipped_an_item_and_need_it_near_the_top_of_my_list_in_libreoffice_writer_how_can_i.docx"
    verify_task(FILE_PATH)
