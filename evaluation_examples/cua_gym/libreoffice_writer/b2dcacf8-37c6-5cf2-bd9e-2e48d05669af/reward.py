"""
FINAL REWARD SCRIPT - SUCCESS
Task: I’ve got my first table ready—the one listing all 152 survey participants—but the document still looks incomplete without a proper label. How do I get LibreOffice Writer to add the exact caption “Table 1: Participants” directly above that table so it’s recognized for the table-of-figures later on?
Generated: 2025-09-10 15:27:19
Status: success
Model: azure-o3
Total Steps: 3
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as _Table
from docx.text.paragraph import Paragraph as _Paragraph


def _iter_block_items(parent):
    """Yield paragraphs and tables in document order.
    This low-level iteration lets us detect exactly which paragraph
    precedes the first table (needed to confirm the caption’s position).
    """
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield _Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield _Table(child, parent)


def verify_caption_for_first_table(file_path: str) -> float:
    """Verify that the document contains a caption “Table 1: Participants”
    directly above the first table, using a caption paragraph style.

    Scoring (progressive):
      • 0.3 – document contains at least one table
      • 0.2 – a paragraph exists immediately above that table
      • 0.3 – paragraph text matches exactly ‘Table 1: Participants’
      • 0.2 – paragraph style name contains “caption” (Writer/Word standard)
    Returns a float between 0.0 and 1.0 and prints detailed diagnostics.
    """
    print(f"Starting verification for file: {file_path}\n")
    score = 0.0
    max_score = 1.0

    # ---------- Prerequisite: file must load ----------
    if not os.path.exists(file_path):
        print("✗ File not found.")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Failed to open document: {exc}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- Requirement 1: table present (0.3) ----------
    if doc.tables:
        print(f"✓ Found {len(doc.tables)} table(s)")
        score += 0.3
    else:
        print("✗ No tables present – caption cannot be verified")
        print(f"REWARD: {score}")
        return score

    # Locate first table in the flow
    blocks = list(_iter_block_items(doc))
    first_table_idx = next((i for i, blk in enumerate(blocks) if isinstance(blk, _Table)), None)
    if first_table_idx is None:
        print("✗ Unexpected: could not locate table among block items")
        print(f"REWARD: {score}")
        return score

    # ---------- Requirement 2: paragraph immediately above table (0.2) ----------
    if first_table_idx == 0:
        print("✗ Table is the first element – no paragraph above it")
    else:
        prev_block = blocks[first_table_idx - 1]
        if isinstance(prev_block, _Paragraph):
            caption_text = prev_block.text.strip()
            print(f"Paragraph above table: '{caption_text}'")
            if caption_text:
                print("✓ Non-empty paragraph located directly above the table")
                score += 0.2

                # ---------- Requirement 3: correct caption text (0.3) ----------
                normalized = re.sub(r"\s+", " ", caption_text)
                pattern = r"^Table\s*1\s*[:\.\-]?\s*Participants$"
                if re.match(pattern, normalized, flags=re.IGNORECASE):
                    print("✓ Caption text matches 'Table 1: Participants'")
                    score += 0.3
                else:
                    print("✗ Caption text does NOT match 'Table 1: Participants'")

                # ---------- Requirement 4: style contains ‘caption’ (0.2) ----------
                style_name = prev_block.style.name if prev_block.style else "(no style)"
                print(f"Paragraph style: {style_name}")
                if "caption" in style_name.lower():
                    print("✓ Paragraph style indicates caption")
                    score += 0.2
                else:
                    print("✗ Paragraph style is not a caption style")
            else:
                print("✗ Paragraph above table is empty – no caption text")
        else:
            print("✗ Element above table is not a paragraph – caption missing")

    # ---------- Final reporting ----------
    final_score = min(score, max_score)
    print("\nVerification complete.")
    print(f"Total score: {final_score}/{max_score}")
    print(f"REWARD: {final_score}")
    return final_score

# ---------------- Execute when run as script ----------------
if __name__ == "__main__":
    test_path = "/home/user/ive_got_my_first_table_readythe_one_listing_all_152_survey_participantsbut_the_document_still_looks_.docx"
    verify_caption_for_first_table(test_path)

