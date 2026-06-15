"""
Reward Script: Insert Dr. Chen citation data as table into 'Impact Metrics' section
Task ID: osworld_multi_apps_calc_to_writer_015
Domain: libreoffice_writer (multi-app: calc + writer)

Task: From ~/research/citation_counts.xlsx, find all papers by 'Dr. Chen'
      and insert their citation data as a table into the 'Impact Metrics' section
      of the research summary document (research_summary.docx / task docx).

Scoring Rubric:
  Component 1: Table exists in document                  (0.3 pts)
  Component 2: Table is placed in 'Impact Metrics' section (0.2 pts)
  Component 3: Table has correct structure (7 rows x 5 cols, correct headers) (0.2 pts)
  Component 4: All 6 Dr. Chen rows are present with correct data              (0.3 pts)
  Total: 1.0
"""

import os

# python-docx for Writer verification
try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("CRITICAL: python-docx not available")
    print("REWARD: 0.0")
    raise SystemExit

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_015'

# Ground truth: 6 Dr. Chen papers from citation_counts.xlsx
DR_CHEN_PAPERS = [
    ('Dr. Chen', 'Deep Learning Approaches for Protein Structure Prediction', 2021, 312, 'Nature Methods'),
    ('Dr. Chen', 'Transformer Architectures in Genomic Sequence Analysis', 2022, 245, 'Bioinformatics'),
    ('Dr. Chen', 'Graph Neural Networks for Molecular Property Prediction', 2020, 189, 'Journal of Chemical Theory'),
    ('Dr. Chen', 'Attention Mechanisms for Single-Cell RNA Sequencing', 2023, 134, 'Cell Systems'),
    ('Dr. Chen', 'Self-Supervised Pretraining on Clinical Text Corpora', 2022, 278, 'npj Digital Medicine'),
    ('Dr. Chen', 'Variational Autoencoders for Drug Candidate Generation', 2020, 196, 'ACS Central Science'),
]

EXPECTED_HEADERS = ['Author', 'Paper Title', 'Year', 'Citations', 'Journal']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists in document (0.3 points)
    # Initial doc has 0 tables; golden doc should have 1 table
    try:
        num_tables = len(doc.tables)
        if num_tables >= 1:
            print(f"PASS: Component 1 — Document contains {num_tables} table(s) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No tables found in document (expected at least 1)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    if total_score < 0.3:
        # No table at all — skip remaining checks
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Table is placed in 'Impact Metrics' section (0.2 points)
    # Verify by checking the table appears after the "4. Impact Metrics" heading
    # in the body element order.
    try:
        body = doc.element.body
        impact_metrics_idx = None
        table_idx = None

        for idx, child in enumerate(body):
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                # Get all text in this paragraph
                text_parts = child.findall('.//' + qn('w:t'))
                para_text = ''.join(t.text or '' for t in text_parts)
                if 'Impact Metrics' in para_text:
                    impact_metrics_idx = idx
            elif tag == 'tbl':
                # Record position of the first table
                if table_idx is None:
                    table_idx = idx

        if impact_metrics_idx is not None and table_idx is not None and table_idx > impact_metrics_idx:
            print(f"PASS: Component 2 — Table (body_idx={table_idx}) is placed after 'Impact Metrics' heading (body_idx={impact_metrics_idx}) (0.2 pts)")
            total_score += 0.2
        else:
            if impact_metrics_idx is None:
                print("FAIL: Component 2 — 'Impact Metrics' section heading not found in document")
            elif table_idx is None:
                print("FAIL: Component 2 — No table found in document body")
            else:
                print(f"FAIL: Component 2 — Table (body_idx={table_idx}) is NOT after 'Impact Metrics' (body_idx={impact_metrics_idx})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table has correct structure (7 rows x 5 cols, correct headers) (0.2 points)
    # Expected: header row + 6 Dr. Chen rows = 7 rows, 5 columns
    try:
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)

        # Check headers in row 0
        header_row = [cell.text.strip() for cell in table.rows[0].cells]
        headers_match = (header_row == EXPECTED_HEADERS)

        if num_rows == 7 and num_cols == 5 and headers_match:
            print(f"PASS: Component 3 — Table has correct structure: {num_rows} rows x {num_cols} cols with headers {header_row} (0.2 pts)")
            total_score += 0.2
        else:
            issues = []
            if num_rows != 7:
                issues.append(f"rows={num_rows} (expected 7)")
            if num_cols != 5:
                issues.append(f"cols={num_cols} (expected 5)")
            if not headers_match:
                issues.append(f"headers={header_row} (expected {EXPECTED_HEADERS})")
            print(f"FAIL: Component 3 — Table structure incorrect: {'; '.join(issues)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: All 6 Dr. Chen rows are present with correct data (0.3 points)
    # Check that each expected Dr. Chen paper appears as a table row
    try:
        table = doc.tables[0]
        # Collect data rows (skip header row)
        table_data_rows = []
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            table_data_rows.append(cells)

        matched_papers = 0
        for paper in DR_CHEN_PAPERS:
            author, title, year, citations, journal = paper
            # Match by author + title (most reliable identifiers)
            found = False
            for row_cells in table_data_rows:
                if (len(row_cells) >= 5 and
                        row_cells[0] == author and
                        row_cells[1] == title):
                    found = True
                    # Optionally verify year, citations, journal
                    year_ok = (str(row_cells[2]) == str(year))
                    citations_ok = (str(row_cells[3]) == str(citations))
                    journal_ok = (row_cells[4] == journal)
                    if year_ok and citations_ok and journal_ok:
                        matched_papers += 1
                        print(f"  MATCH: '{title[:40]}...' year={year} citations={citations}")
                    else:
                        issues = []
                        if not year_ok:
                            issues.append(f"year={row_cells[2]} (expected {year})")
                        if not citations_ok:
                            issues.append(f"citations={row_cells[3]} (expected {citations})")
                        if not journal_ok:
                            issues.append(f"journal={row_cells[4]!r} (expected {journal!r})")
                        print(f"  PARTIAL: '{title[:40]}...' data mismatch: {'; '.join(issues)}")
                    break
            if not found:
                print(f"  MISSING: '{title[:50]}...'")

        if matched_papers == 6:
            print(f"PASS: Component 4 — All 6 Dr. Chen papers found with correct data (0.3 pts)")
            total_score += 0.3
        elif matched_papers > 0:
            # Partial credit proportional to number matched
            partial = round(0.3 * matched_papers / 6, 4)
            print(f"PARTIAL: Component 4 — {matched_papers}/6 Dr. Chen papers found with correct data ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No Dr. Chen papers found in table data rows")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
