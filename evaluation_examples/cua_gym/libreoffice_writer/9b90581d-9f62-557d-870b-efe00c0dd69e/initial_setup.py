"""
Initial Setup: Multi-app task — citation data in calc + research summary in writer
Task ID: osworld_multi_apps_calc_to_writer_015
Domain: libreoffice_writer (multi-app: calc + writer)

Creates:
  - ~/research/citation_counts.xlsx  (spreadsheet with citation data including Dr. Chen rows)
  - ~/osworld_multi_apps_calc_to_writer_015.docx  (research summary with Impact Metrics section, no table yet)

Then opens both files: the .docx in LibreOffice Writer (primary) and the .xlsx in LibreOffice Calc.
"""

import os
import shlex
import subprocess
import time

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_015'
RESEARCH_DIR = f'{WORKDIR}/research'
XLSX_PATH = f'{RESEARCH_DIR}/citation_counts.xlsx'
DOCX_PATH = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_citation_spreadsheet():
    """Create ~/research/citation_counts.xlsx with realistic citation data."""
    os.makedirs(RESEARCH_DIR, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Citations"

    # Header row
    headers = ["Author", "Paper Title", "Year", "Citations", "Journal"]
    header_font = Font(name="Calibri", bold=True, size=12)
    header_fill = PatternFill(start_color="FF4472C4", end_color="FF4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    # Column widths
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 52
    ws.column_dimensions["C"].width = 8
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 30
    ws.row_dimensions[1].height = 22

    # Data rows — realistic multi-author citation data
    # Dr. Chen rows will be rows that agent needs to extract
    data = [
        ["Dr. Chen", "Deep Learning Approaches for Protein Structure Prediction", 2021, 312, "Nature Methods"],
        ["Prof. Martinez", "Statistical Models in Epidemiological Surveillance", 2020, 178, "Epidemiology & Infection"],
        ["Dr. Chen", "Transformer Architectures in Genomic Sequence Analysis", 2022, 245, "Bioinformatics"],
        ["Dr. Patel", "Reinforcement Learning for Autonomous Drug Discovery", 2021, 203, "Journal of Chemical Information"],
        ["Dr. Chen", "Graph Neural Networks for Molecular Property Prediction", 2020, 189, "Journal of Chemical Theory"],
        ["Prof. Thompson", "Multi-Omics Integration in Cancer Biomarker Discovery", 2022, 156, "Cancer Research"],
        ["Dr. Lee", "Federated Learning in Clinical Trial Data Sharing", 2023, 98, "JAMIA"],
        ["Dr. Chen", "Attention Mechanisms for Single-Cell RNA Sequencing", 2023, 134, "Cell Systems"],
        ["Dr. Nguyen", "Bayesian Inference for Causal Modeling in Genomics", 2021, 167, "PLOS Genetics"],
        ["Prof. Williams", "Contrastive Learning in Multimodal Medical Imaging", 2022, 211, "Medical Image Analysis"],
        ["Dr. Chen", "Self-Supervised Pretraining on Clinical Text Corpora", 2022, 278, "npj Digital Medicine"],
        ["Dr. Kim", "Large Language Models for Electronic Health Records", 2023, 321, "Journal of Biomedical Informatics"],
        ["Prof. Anderson", "CRISPR-Cas9 Off-Target Effect Prediction via Deep Learning", 2021, 145, "Nucleic Acids Research"],
        ["Dr. Chen", "Variational Autoencoders for Drug Candidate Generation", 2020, 196, "ACS Central Science"],
        ["Dr. Okafor", "Time-Series Analysis of Wearable Biosensor Data", 2023, 87, "NPJ Digital Medicine"],
    ]

    thin_border = Border(
        left=Side(style="thin", color="AAAAAA"),
        right=Side(style="thin", color="AAAAAA"),
        top=Side(style="thin", color="AAAAAA"),
        bottom=Side(style="thin", color="AAAAAA"),
    )

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = Font(name="Calibri", size=11)
            cell.border = thin_border
            if c == 4:  # Citations column — right-align numbers
                cell.alignment = Alignment(horizontal="right")
            if c == 3:  # Year column — center
                cell.alignment = Alignment(horizontal="center")
        # Zebra striping for readability
        if r % 2 == 0:
            for c in range(1, 6):
                ws.cell(row=r, column=c).fill = PatternFill(
                    start_color="FFF2F2F2", end_color="FFF2F2F2", fill_type="solid"
                )

    wb.save(XLSX_PATH)
    print(f"Citation spreadsheet created: {XLSX_PATH}")


def create_research_summary():
    """Create the research summary .docx with 'Impact Metrics' section but NO table."""
    doc = Document()

    # ---- Title ----
    title = doc.add_heading("Computational Biology Research Summary", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Subtitle / byline ----
    byline = doc.add_paragraph("Prepared by the Institute for Computational Life Sciences")
    byline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    byline.paragraph_format.space_after = Pt(12)
    for run in byline.runs:
        run.font.italic = True
        run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # ---- Section 1: Overview ----
    doc.add_heading("1. Overview", level=1)
    overview_text = (
        "This document summarizes the research output of our institute's computational biology "
        "division for the period 2020–2023. The division focuses on applying machine learning "
        "and statistical methods to problems in genomics, drug discovery, and clinical informatics. "
        "The following sections provide a narrative summary, key collaborations, and bibliometric "
        "impact data for selected investigators."
    )
    doc.add_paragraph(overview_text)

    # ---- Section 2: Research Focus Areas ----
    doc.add_heading("2. Research Focus Areas", level=1)
    focus_intro = doc.add_paragraph(
        "The division's work spans three primary research thrusts:"
    )

    doc.add_paragraph(
        "Structural and Molecular Biology: Computational prediction of protein structures, "
        "molecular docking, and small-molecule property estimation using deep neural architectures.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Genomics and Transcriptomics: Single-cell RNA sequencing analysis, genomic sequence "
        "modeling, and multi-omics data integration for biomarker identification.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Clinical and Translational Informatics: Natural language processing on clinical text, "
        "electronic health record mining, and federated learning frameworks for privacy-preserving "
        "data sharing across hospital networks.",
        style="List Bullet"
    )

    doc.add_paragraph(
        "Each thrust has produced high-impact publications in leading journals and contributed "
        "to collaborative grants totaling over $12 million in external funding over the review period."
    )

    # ---- Section 3: Key Collaborations ----
    doc.add_heading("3. Key Collaborations", level=1)
    collab_para = doc.add_paragraph(
        "Our investigators maintain active collaborations with ten partner institutions across "
        "North America, Europe, and Asia. Notable partnerships include joint NIH R01 grants with "
        "Stanford University and the Broad Institute, an EU Horizon project on federated clinical "
        "data analysis, and an industry collaboration with a leading pharmaceutical company for "
        "AI-driven compound screening. These relationships have enhanced both the scientific quality "
        "and translational potential of our outputs."
    )

    # ---- Section 4: Impact Metrics ---- (KEY SECTION — NO TABLE HERE YET)
    doc.add_heading("4. Impact Metrics", level=1)

    impact_intro = doc.add_paragraph(
        "The table below presents citation data for selected publications. The data is drawn "
        "from the institute's citation tracking database and covers peer-reviewed articles "
        "published between 2020 and 2023."
    )
    impact_intro.paragraph_format.space_after = Pt(8)

    # NOTE: No table is inserted here. The agent must insert the Dr. Chen citation table.
    # Placeholder paragraph so the section is non-empty but clearly lacks the table.
    placeholder = doc.add_paragraph(
        "[Citation table to be inserted here]"
    )
    for run in placeholder.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    placeholder.paragraph_format.space_after = Pt(10)

    # ---- Section 5: Conclusions ----
    doc.add_heading("5. Conclusions", level=1)
    conclusions = doc.add_paragraph(
        "The division's output over the 2020–2023 period demonstrates sustained high impact "
        "across all three focus areas. Citation metrics indicate strong community uptake of the "
        "published methodologies. Looking ahead, leadership has approved expansion of the clinical "
        "informatics team and additional computational infrastructure investment to support "
        "large-scale foundation model training for biomedical applications."
    )

    doc.save(DOCX_PATH)
    print(f"Research summary document created: {DOCX_PATH}")


def main():
    create_citation_spreadsheet()
    create_research_summary()

    # GUI-ready startup: open Writer with the research summary (primary task file)
    # Also open Calc with citation_counts.xlsx so agent can reference it
    launch_gui(f'libreoffice --writer "{DOCX_PATH}"', delay_sec=2.5)
    launch_gui(f'libreoffice --calc "{XLSX_PATH}"', delay_sec=2.0)

    print("GUI_READY: launched LibreOffice Writer (research_summary) and Calc (citation_counts) with DISPLAY=:0")


main()
