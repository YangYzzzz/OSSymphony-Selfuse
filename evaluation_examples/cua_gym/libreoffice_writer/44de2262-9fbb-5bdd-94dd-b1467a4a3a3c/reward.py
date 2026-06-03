"""
Reward Script: Add APA bibliography entry and cross-reference citation
Task ID: osworld_writer_bibliography_crossref_009
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Martinez & Park (2022) bibliography entry present in References section
                     in alphabetical order (between Chen and Thompson entries)
  Component 2 (0.5): Citation '(Martinez & Park, 2022)' inserted in the last sentence
                     of the 4th paragraph of the Introduction section
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_009'

# Expected bibliography entry (key fragments)
EXPECTED_ENTRY_AUTHOR = 'Martinez, S., & Park, J. (2022)'
EXPECTED_ENTRY_TITLE = 'International Trade Agreements: A Comprehensive Review'
EXPECTED_ENTRY_PUBLISHER = 'Global Economics Press'

# Expected citation in body text
EXPECTED_CITATION = '(Martinez & Park, 2022)'

# In initial doc, References section has 4 entries (Anderson, Bhagwati, Chen, Thompson)
# The 4th paragraph of Introduction is index 11 (0-based) in the document:
#   Para 7 = Heading "Introduction"
#   Para 8 = 1st body paragraph
#   Para 9 = 2nd body paragraph
#   Para 10 = 3rd body paragraph
#   Para 11 = 4th body paragraph  <-- citation goes here


def find_introduction_paragraphs(doc):
    """Return a list of (index, paragraph) tuples for body paragraphs under Introduction."""
    intro_paras = []
    in_intro = False
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 1' and para.text.strip() == 'Introduction':
            in_intro = True
            continue
        if in_intro:
            if para.style.name == 'Heading 1':
                # Hit the next section heading — stop
                break
            intro_paras.append((i, para))
    return intro_paras


def find_references_paragraphs(doc):
    """Return a list of (index, paragraph) tuples for entries in References section."""
    ref_paras = []
    in_refs = False
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 1' and para.text.strip() == 'References':
            in_refs = True
            continue
        if in_refs:
            if para.style.name == 'Heading 1':
                break
            if para.text.strip():
                ref_paras.append((i, para))
    return ref_paras


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Martinez & Park (2022) bibliography entry in References (0.5 pts)
    # The entry must exist and must be in alphabetical order:
    # alphabetically between Chen (C) and Thompson (T), i.e., at position after Chen
    # -------------------------------------------------------------------------
    try:
        ref_paras = find_references_paragraphs(doc)
        ref_texts = [para.text.strip() for (_, para) in ref_paras]

        # Check whether the Martinez entry is present with key fragments
        martinez_entry_found = any(
            EXPECTED_ENTRY_AUTHOR in t and EXPECTED_ENTRY_TITLE in t and EXPECTED_ENTRY_PUBLISHER in t
            for t in ref_texts
        )

        if not martinez_entry_found:
            print(f"FAIL: Component 1 — Martinez & Park (2022) entry not found in References section.")
            print(f"      Current entries: {ref_texts}")
        else:
            # Check alphabetical ordering: Martinez should come after Chen and before Thompson
            martinez_idx = next(
                (i for i, t in enumerate(ref_texts) if EXPECTED_ENTRY_AUTHOR in t), None
            )
            chen_idx = next(
                (i for i, t in enumerate(ref_texts) if t.startswith('Chen,')), None
            )
            thompson_idx = next(
                (i for i, t in enumerate(ref_texts) if t.startswith('Thompson,')), None
            )

            order_correct = (
                chen_idx is not None
                and thompson_idx is not None
                and martinez_idx is not None
                and chen_idx < martinez_idx < thompson_idx
            )

            if order_correct:
                print(f"PASS: Component 1 — Martinez entry found in References in correct alphabetical order "
                      f"(position {martinez_idx} among {len(ref_texts)} entries, after Chen at {chen_idx}, "
                      f"before Thompson at {thompson_idx}) (0.5 pts)")
                total_score += 0.5
            else:
                # Entry exists but not in correct alphabetical position — partial check
                print(f"FAIL: Component 1 — Martinez entry found but NOT in correct alphabetical order. "
                      f"chen_idx={chen_idx}, martinez_idx={martinez_idx}, thompson_idx={thompson_idx}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Citation '(Martinez & Park, 2022)' in last sentence of 4th paragraph (0.5 pts)
    # The 4th body paragraph of the Introduction section (0-based: 3rd in the list
    # of body paragraphs under the Introduction heading).
    # -------------------------------------------------------------------------
    try:
        intro_body_paras = find_introduction_paragraphs(doc)
        # Filter out blank paragraphs to get actual body text paragraphs
        nonempty_intro = [(i, p) for (i, p) in intro_body_paras if p.text.strip()]

        if len(nonempty_intro) < 4:
            print(f"FAIL: Component 2 — Expected at least 4 paragraphs in Introduction, "
                  f"found {len(nonempty_intro)}")
        else:
            # The 4th body paragraph (index 3 in 0-based list)
            fourth_para_idx, fourth_para = nonempty_intro[3]
            para_text = fourth_para.text

            if EXPECTED_CITATION in para_text:
                print(f"PASS: Component 2 — Citation '{EXPECTED_CITATION}' found in the 4th paragraph "
                      f"of Introduction (doc para index {fourth_para_idx}) (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 2 — Citation '{EXPECTED_CITATION}' NOT found in the 4th "
                      f"paragraph of Introduction (doc para index {fourth_para_idx}).")
                print(f"      Paragraph ends with: ...{repr(para_text[-150:])}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
