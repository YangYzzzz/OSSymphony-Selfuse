"""
Initial Setup: Add APA-format bibliography entry and cross-reference in economics research paper
Task ID: osworld_writer_bibliography_crossref_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading('Global Trade Policy and Economic Development', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / author info
    author_para = doc.add_paragraph('Elena Vasquez and James O\'Brien')
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.runs[0].font.size = Pt(12)

    affil_para = doc.add_paragraph('Department of Economics, Eastfield University')
    affil_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil_para.runs[0].font.size = Pt(11)

    doc.add_paragraph('')  # blank line

    # Abstract heading
    abstract_heading = doc.add_heading('Abstract', level=1)

    abstract_text = (
        'This paper examines the multifaceted relationship between global trade policy '
        'and economic development in emerging markets. Drawing on empirical data from '
        '37 countries over a 20-year period, we analyze how tariff structures, trade '
        'agreements, and regulatory frameworks shape growth trajectories. Our findings '
        'suggest that balanced liberalization paired with institutional capacity building '
        'yields the most sustainable development outcomes.'
    )
    doc.add_paragraph(abstract_text)

    doc.add_paragraph('')

    # Body heading
    doc.add_heading('Introduction', level=1)

    # Paragraph 1
    para1_text = (
        'The relationship between international trade and economic development has been '
        'a central concern of economists since the foundational work of Ricardo (1817) '
        'and later elaborated by Heckscher and Ohlin in the early twentieth century. '
        'Over the past three decades, the pace of trade liberalization has accelerated '
        'dramatically, driven in part by the establishment of the World Trade Organization '
        'in 1995 and subsequent rounds of multilateral negotiations. Scholars such as '
        'Bhagwati (2004) and Rodrik (2011) have offered contrasting perspectives on '
        'whether these developments have broadly benefited developing economies or '
        'exacerbated pre-existing inequalities.'
    )
    doc.add_paragraph(para1_text)

    # Paragraph 2
    para2_text = (
        'Regional trade agreements have proliferated in tandem with multilateral efforts, '
        'creating a complex web of overlapping commitments and preferential arrangements. '
        'The number of such agreements in force exceeded 350 by 2020, covering more than '
        'half of world merchandise trade (Anderson & Kee, 2019). While proponents argue '
        'that regional integration fosters economies of scale and technology transfer, '
        'critics caution that rules-of-origin provisions and trade diversion effects '
        'may offset these gains, particularly for smaller or less competitive economies '
        'that lack bargaining leverage in bilateral negotiations.'
    )
    doc.add_paragraph(para2_text)

    # Paragraph 3
    para3_text = (
        'Empirical assessments of trade policy impacts have increasingly relied on '
        'sophisticated econometric techniques, including difference-in-differences '
        'estimation and synthetic control methods, to identify causal effects rather '
        'than mere correlations (Chen & Liu, 2020). These methodological advances have '
        'helped clarify contested questions about the distributional consequences of '
        'openness, revealing that aggregate gains often mask significant heterogeneity '
        'across sectors and income groups within countries. The findings underscore '
        'the importance of complementary policies in labor markets, education, and '
        'social protection to ensure that the benefits of trade are broadly shared.'
    )
    doc.add_paragraph(para3_text)

    # Paragraph 4 — the citation needs to go at the end of the last sentence here
    # NOTE: NO "(Martinez & Park, 2022)" citation yet — that is the task
    para4_text = (
        'The governance architecture of international trade has also evolved considerably, '
        'with dispute settlement mechanisms playing an increasingly prominent role in '
        'adjudicating conflicts between trading partners. The WTO Appellate Body, despite '
        'facing operational challenges in recent years, remains a critical forum for '
        'resolving disagreements over subsidy practices, anti-dumping measures, and '
        'sanitary regulations. Beyond formal dispute resolution, bilateral investment '
        'treaties and investor-state arbitration clauses have attracted scrutiny for '
        'potentially constraining the policy space of national governments in managing '
        'their development strategies. The normative framework governing trade has '
        'therefore become inseparable from broader questions about sovereignty, '
        'democratic accountability, and the equitable distribution of gains from '
        'economic integration.'
    )
    doc.add_paragraph(para4_text)

    # Paragraph 5
    para5_text = (
        'Looking ahead, the trajectory of global trade policy will be shaped by several '
        'converging forces: the digital economy, climate change imperatives, and '
        'geopolitical realignments triggered by shifting power dynamics among major '
        'trading nations. The rise of e-commerce and data-driven services has outpaced '
        'existing regulatory frameworks, creating governance gaps that existing trade '
        'law was not designed to address (Thompson & Williams, 2021). At the same time, '
        'the imperative to decarbonize industrial supply chains is prompting fresh '
        'debates about carbon border adjustments and environmental conditionality in '
        'trade agreements. How policymakers navigate these complex trade-offs will '
        'determine whether international trade continues to serve as an engine of '
        'inclusive growth or becomes a source of renewed fragmentation and conflict.'
    )
    doc.add_paragraph(para5_text)

    doc.add_paragraph('')

    # References section
    ref_heading = doc.add_heading('References', level=1)

    # 4 existing reference entries in alphabetical order
    # Note: Martinez & Park (2022) is NOT included — that is the task
    refs = [
        'Anderson, J., & Kee, H. L. (2019). Regional trade agreements and global welfare. '
        'Journal of International Economics, 114, 78–95.',

        'Bhagwati, J. (2004). In defense of globalization. Oxford University Press, New York.',

        'Chen, Y., & Liu, M. (2020). Causal identification in trade policy research: Methods '
        'and applications. Annual Review of Economics, 12, 345–378.',

        'Thompson, R., & Williams, A. (2021). Digital trade governance in the post-pandemic '
        'era. World Economy, 44(8), 2301–2325.',
    ]

    for ref in refs:
        ref_para = doc.add_paragraph(ref)
        # Hanging indent for APA style
        ref_para.paragraph_format.left_indent = Pt(36)
        ref_para.paragraph_format.first_line_indent = Pt(-36)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
