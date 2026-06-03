"""
Initial Setup: Multi-app Writer+GIMP photo processing workflow task
Task ID: osworld_multi_apps_writer_gimp_061
Domain: libreoffice_writer + gimp

Creates:
  - /home/user/Desktop/workflow.docx: Instructions doc with 3 processing steps
  - /home/user/Desktop/raw_photo.jpg: Realistic-looking portrait photo
Opens:
  - LibreOffice Writer with workflow.docx
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from PIL import Image, ImageDraw, ImageFilter
import numpy as np

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_gimp_061'
DOC_PATH = f'{WORKDIR}/workflow.docx'
PHOTO_PATH = f'{WORKDIR}/raw_photo.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_workflow_doc():
    """Create workflow.docx with three photo processing steps."""
    doc = Document()

    # Title
    title = doc.add_heading('Photo Processing Workflow', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    intro = doc.add_paragraph(
        'This document outlines the multi-step processing pipeline for the portrait photo. '
        'Follow each step in order using GIMP to achieve professional-quality results. '
        'Apply all steps to "raw_photo.jpg" and save the final result as "processed_photo.jpg" on the Desktop.'
    )
    intro.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # blank line

    # Step 1 heading
    step1_heading = doc.add_heading('Step 1: Noise Reduction', level=1)

    # Step 1 instructions
    step1_para = doc.add_paragraph()
    run = step1_para.add_run('Apply Gaussian blur with radius ')
    run.font.size = Pt(11)
    run_val = step1_para.add_run('1.0')
    run_val.bold = True
    run_val.font.size = Pt(11)
    run_val.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
    run2 = step1_para.add_run(' to reduce noise across the entire image.')
    run2.font.size = Pt(11)

    detail1 = doc.add_paragraph(
        'In GIMP: Filters > Blur > Gaussian Blur. '
        'Set Size X and Size Y both to 1.0 pixels. Apply to the flattened image.'
    )
    detail1.paragraph_format.left_indent = Pt(24)
    detail1.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # blank line

    # Step 2 heading
    step2_heading = doc.add_heading('Step 2: Skin Tone Correction', level=1)

    # Step 2 instructions
    step2_para = doc.add_paragraph()
    run3 = step2_para.add_run('Adjust hue/saturation to correct skin tones in the red/orange range:')
    run3.font.size = Pt(11)

    # Bullet list for step 2 params
    bullet1 = doc.add_paragraph(style='List Bullet')
    r_b1 = bullet1.add_run('Hue shift: ')
    r_b1.font.size = Pt(11)
    r_b1_val = bullet1.add_run('+3')
    r_b1_val.bold = True
    r_b1_val.font.size = Pt(11)
    r_b1_val.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
    r_b1_rest = bullet1.add_run(' degrees (red/orange channel)')
    r_b1_rest.font.size = Pt(11)

    bullet2 = doc.add_paragraph(style='List Bullet')
    r_b2 = bullet2.add_run('Saturation adjustment: ')
    r_b2.font.size = Pt(11)
    r_b2_val = bullet2.add_run('-5')
    r_b2_val.bold = True
    r_b2_val.font.size = Pt(11)
    r_b2_val.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
    r_b2_rest = bullet2.add_run(' (reduce saturation to natural skin tones)')
    r_b2_rest.font.size = Pt(11)

    detail2 = doc.add_paragraph(
        'In GIMP: Colors > Hue-Saturation. Select the R (Red) and/or Y (Yellow-Orange) ranges. '
        'Set Hue to +3 and Saturation to -5. Click OK.'
    )
    detail2.paragraph_format.left_indent = Pt(24)
    detail2.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # blank line

    # Step 3 heading
    step3_heading = doc.add_heading('Step 3: Background Blur', level=1)

    # Step 3 instructions
    step3_para = doc.add_paragraph()
    run4 = step3_para.add_run('Apply Gaussian blur with radius ')
    run4.font.size = Pt(11)
    run4_val = step3_para.add_run('8.0')
    run4_val.bold = True
    run4_val.font.size = Pt(11)
    run4_val.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
    run5 = step3_para.add_run(' to the background using a layer mask to simulate depth of field.')
    run5.font.size = Pt(11)

    detail3 = doc.add_paragraph(
        'In GIMP: Duplicate the layer. Add a layer mask to the top layer (white). '
        'Paint the subject area black on the mask to protect it. '
        'Apply Filters > Blur > Gaussian Blur with radius 8.0 to the background layer.'
    )
    detail3.paragraph_format.left_indent = Pt(24)
    detail3.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # blank line

    # Summary / footer note
    note = doc.add_paragraph(
        'Note: After applying all three steps, flatten the image (Image > Flatten Image) '
        'and export as "processed_photo.jpg" to the Desktop using File > Export As.'
    )
    run_note = note.runs[0]
    run_note.font.italic = True
    run_note.font.size = Pt(10)
    run_note.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Ensure Desktop dir exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(DOC_PATH)
    print(f'Workflow document created: {DOC_PATH}')


def create_raw_photo():
    """Create a realistic portrait-style photo as raw_photo.jpg."""
    width, height = 800, 1000
    np.random.seed(42)

    # Build a realistic-looking portrait scene using Pillow
    img = Image.new('RGB', (width, height), color=(180, 200, 220))
    draw = ImageDraw.Draw(img)

    # --- Background: gradient sky + bokeh hint ---
    bg_arr = np.zeros((height, width, 3), dtype=np.uint8)
    for y in range(height):
        t = y / height
        r = int(140 + t * 60)
        g = int(160 + t * 40)
        b = int(200 - t * 30)
        bg_arr[y, :] = [r, g, b]
    img = Image.fromarray(bg_arr)
    draw = ImageDraw.Draw(img)

    # Bokeh-like soft blobs in background
    rng = np.random.RandomState(123)
    for _ in range(20):
        bx = rng.randint(0, width)
        by = rng.randint(0, int(height * 0.7))
        br = rng.randint(15, 50)
        bc = (rng.randint(200, 255), rng.randint(200, 255), rng.randint(200, 255))
        draw.ellipse([bx - br, by - br, bx + br, by + br], fill=bc)

    # Blur the background slightly
    img = img.filter(ImageFilter.GaussianBlur(radius=4))
    draw = ImageDraw.Draw(img)

    # --- Subject: person silhouette (torso + head) ---
    cx = width // 2

    # Torso/shoulders
    torso_top = int(height * 0.42)
    torso_bot = height + 10
    shoulder_w = int(width * 0.38)
    # Body fill - slightly warm skin/clothing tone
    draw.ellipse([cx - shoulder_w, torso_top, cx + shoulder_w, torso_bot],
                 fill=(90, 70, 55))
    # Shirt/clothing
    draw.rectangle([cx - shoulder_w + 10, torso_top + 80, cx + shoulder_w - 10, torso_bot],
                   fill=(45, 85, 130))

    # Neck
    neck_w = 38
    neck_top = int(height * 0.35)
    neck_bot = torso_top + 20
    draw.rectangle([cx - neck_w, neck_top, cx + neck_w, neck_bot], fill=(210, 170, 140))

    # Head - oval shape
    head_cx = cx
    head_cy = int(height * 0.27)
    head_w = 120
    head_h = 150
    draw.ellipse([head_cx - head_w, head_cy - head_h,
                  head_cx + head_w, head_cy + head_h],
                 fill=(220, 180, 150))

    # Hair
    hair_cy = head_cy - head_h + 30
    draw.ellipse([head_cx - head_w - 5, hair_cy - 50,
                  head_cx + head_w + 5, head_cy + 60],
                 fill=(55, 35, 20))
    draw.rectangle([head_cx - head_w - 5, hair_cy,
                    head_cx + head_w + 5, hair_cy + 30],
                   fill=(55, 35, 20))

    # Eyes
    eye_y = head_cy - 10
    for ex in [head_cx - 45, head_cx + 45]:
        draw.ellipse([ex - 18, eye_y - 10, ex + 18, eye_y + 10], fill=(255, 255, 240))
        draw.ellipse([ex - 8, eye_y - 7, ex + 8, eye_y + 7], fill=(65, 45, 30))
        draw.ellipse([ex - 4, eye_y - 4, ex + 4, eye_y + 4], fill=(10, 10, 10))

    # Eyebrows
    for ex in [head_cx - 45, head_cx + 45]:
        draw.line([(ex - 20, eye_y - 22), (ex + 20, eye_y - 18)],
                  fill=(45, 30, 15), width=5)

    # Nose
    nose_y = head_cy + 25
    draw.ellipse([head_cx - 15, nose_y - 10, head_cx + 15, nose_y + 15],
                 fill=(200, 160, 130))

    # Mouth
    mouth_y = head_cy + 60
    draw.arc([head_cx - 30, mouth_y - 10, head_cx + 30, mouth_y + 20],
             start=0, end=180, fill=(160, 80, 80), width=4)

    # Add some natural noise to make the photo look more realistic
    arr = np.array(img).astype(np.float32)
    noise = rng.normal(0, 8, arr.shape).astype(np.float32)
    arr = np.clip(arr + noise, 0, 255).astype(np.uint8)
    img = Image.fromarray(arr)

    # Slight overall blur to simulate camera softness (photo-like)
    img = img.filter(ImageFilter.GaussianBlur(radius=0.5))

    os.makedirs(WORKDIR, exist_ok=True)
    img.save(PHOTO_PATH, 'JPEG', quality=92)
    print(f'Raw photo created: {PHOTO_PATH}')


def main():
    create_workflow_doc()
    create_raw_photo()

    # GUI-ready startup: open workflow.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_PATH}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with workflow.docx (DISPLAY=:0)')


main()
