"""
Reward Script: Convert bulleted list to numbered list with 'Step N)' format
Task ID: writer_lec_027
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): List items use decimal numbering (not bullet)
  Component 2 (0.40): Numbering format is 'Step %1)' (custom prefix/suffix)
  Component 3 (0.20): Text content of list items preserved AND numbered
  Component 4 (0.10): Exactly 6 decimal-numbered items
"""

import os
from docx import Document
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_027'

# Expected text content of the 6 list items (unchanged from initial)
EXPECTED_ITEMS = [
    'Export customer records from the legacy CRM database using the Admin Console export wizard',
    'Validate the exported CSV files against the schema definition to ensure all required fields are present',
    'Run the deduplication script to identify and merge duplicate contact entries across regional databases',
    'Transform the cleaned dataset into the target format using the ETL pipeline configuration template',
    'Load the transformed records into the staging environment and execute the automated integration tests',
    'Promote the verified dataset to the production analytics platform and confirm dashboard metrics update correctly',
]

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_para_num_id(para):
    """Get the numId from a paragraph's pPr/numPr, or None."""
    numPr = para._element.find(f'{{{WNS}}}pPr/{{{WNS}}}numPr')
    if numPr is None:
        return None
    numId_elem = numPr.find(f'{{{WNS}}}numId')
    if numId_elem is None:
        return None
    return numId_elem.get(f'{{{WNS}}}val')


def get_numbering_info(doc, num_id):
    """Look up the lvlText and numFmt for a given numId at ilvl=0."""
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return None, None
        numbering_elem = numbering_part.element

        # Find the w:num element with matching numId -> get abstractNumId
        abstract_num_id = None
        for num_elem in numbering_elem.findall(f'{{{WNS}}}num'):
            nid = num_elem.get(f'{{{WNS}}}numId')
            if nid == str(num_id):
                abstract_ref = num_elem.find(f'{{{WNS}}}abstractNumId')
                if abstract_ref is not None:
                    abstract_num_id = abstract_ref.get(f'{{{WNS}}}val')
                break

        if abstract_num_id is None:
            return None, None

        # Find the abstractNum and get lvlText and numFmt for ilvl=0
        for abstract_num in numbering_elem.findall(f'{{{WNS}}}abstractNum'):
            aid = abstract_num.get(f'{{{WNS}}}abstractNumId')
            if aid == abstract_num_id:
                for lvl in abstract_num.findall(f'{{{WNS}}}lvl'):
                    ilvl = lvl.get(f'{{{WNS}}}ilvl')
                    if ilvl == '0':
                        lvl_text_elem = lvl.find(f'{{{WNS}}}lvlText')
                        num_fmt_elem = lvl.find(f'{{{WNS}}}numFmt')
                        lvl_text = lvl_text_elem.get(f'{{{WNS}}}val') if lvl_text_elem is not None else None
                        num_fmt = num_fmt_elem.get(f'{{{WNS}}}val') if num_fmt_elem is not None else None
                        return lvl_text, num_fmt
        return None, None
    except Exception:
        return None, None


def is_decimal_numbered(doc, para):
    """Check if a paragraph uses decimal numbering (not bullet)."""
    num_id = get_para_num_id(para)
    if num_id is None:
        return False
    lvl_text, num_fmt = get_numbering_info(doc, num_id)
    return num_fmt == 'decimal'


def has_step_format(doc, para):
    """Check if a paragraph uses 'Step N)' numbering format."""
    num_id = get_para_num_id(para)
    if num_id is None:
        return False
    lvl_text, num_fmt = get_numbering_info(doc, num_id)
    if lvl_text is None or num_fmt is None:
        return False
    # Must be decimal format with 'Step' prefix and ')' suffix
    has_step = 'step' in lvl_text.lower()
    has_paren = ')' in lvl_text
    is_decimal = num_fmt == 'decimal'
    return has_step and has_paren and is_decimal


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Identify list paragraphs (indices 2-7 in the document)
    if len(doc.paragraphs) < 8:
        print(f"FAIL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 8")
        print("REWARD: 0.0")
        return 0.0

    list_paras = [doc.paragraphs[i] for i in range(2, 8)]

    # Component 1: List items use decimal numbering, not bullet (0.30 points)
    # Initial state has bullet numbering (numFmt=bullet), golden has decimal
    try:
        decimal_count = 0
        for para in list_paras:
            if is_decimal_numbered(doc, para):
                decimal_count += 1

        if decimal_count == 6:
            print(f"PASS: Component 1 — All 6 items use decimal numbering (0.30 pts)")
            total_score += 0.30
        elif decimal_count > 0:
            partial = round(0.30 * (decimal_count / 6), 2)
            print(f"PARTIAL: Component 1 — {decimal_count}/6 items use decimal numbering ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No items use decimal numbering (still bullet)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Numbering format is 'Step %1)' (0.40 points)
    # The key task requirement: custom numbering with 'Step' prefix and ')' suffix
    try:
        step_count = 0
        for para in list_paras:
            if has_step_format(doc, para):
                step_count += 1

        if step_count == 6:
            print(f"PASS: Component 2 — All 6 items use 'Step N)' numbering format (0.40 pts)")
            total_score += 0.40
        elif step_count > 0:
            partial = round(0.40 * (step_count / 6), 2)
            print(f"PARTIAL: Component 2 — {step_count}/6 items use 'Step N)' format ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No items use 'Step N)' numbering format")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text content preserved AND items are decimal-numbered (0.20 points)
    # Gate: only award if items are decimal-numbered (to differentiate from initial state)
    try:
        matched = 0
        for idx, para in enumerate(list_paras):
            if idx < len(EXPECTED_ITEMS):
                actual = para.text.strip()
                expected = EXPECTED_ITEMS[idx].strip()
                if actual == expected:
                    matched += 1
                else:
                    print(f"  INFO: Item {idx+1} text mismatch")

        decimal_count = sum(1 for p in list_paras if is_decimal_numbered(doc, p))

        if matched == 6 and decimal_count == 6:
            print(f"PASS: Component 3 — All 6 item texts preserved and decimal-numbered (0.20 pts)")
            total_score += 0.20
        elif matched == 6 and decimal_count > 0:
            partial = round(0.20 * (decimal_count / 6), 2)
            print(f"PARTIAL: Component 3 — Text preserved, {decimal_count}/6 decimal-numbered ({partial} pts)")
            total_score += partial
        elif matched == 6 and decimal_count == 0:
            print(f"FAIL: Component 3 — Text preserved but no decimal numbering (gate failed)")
        else:
            print(f"FAIL: Component 3 — Text not fully preserved ({matched}/6 match)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Exactly 6 decimal-numbered items (0.10 points)
    try:
        decimal_count = sum(1 for p in list_paras if is_decimal_numbered(doc, p))

        if decimal_count == 6:
            print(f"PASS: Component 4 — Exactly 6 decimal-numbered items found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Found {decimal_count} decimal-numbered items, expected 6")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
