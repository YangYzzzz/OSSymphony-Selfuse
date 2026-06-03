"""
Initial Setup: Create a Writer document with a bulleted list of 6 workflow steps
Task ID: writer_lec_027
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a heading for context
    doc.add_heading('Data Migration Workflow', level=1)

    # Add a brief intro paragraph
    doc.add_paragraph(
        'Follow these steps to complete the quarterly data migration '
        'from the legacy CRM system to the new analytics platform.'
    )

    # 6 bulleted workflow steps (realistic content)
    steps = [
        'Export customer records from the legacy CRM database using the Admin Console export wizard',
        'Validate the exported CSV files against the schema definition to ensure all required fields are present',
        'Run the deduplication script to identify and merge duplicate contact entries across regional databases',
        'Transform the cleaned dataset into the target format using the ETL pipeline configuration template',
        'Load the transformed records into the staging environment and execute the automated integration tests',
        'Promote the verified dataset to the production analytics platform and confirm dashboard metrics update correctly',
    ]

    for step_text in steps:
        doc.add_paragraph(step_text, style='List Bullet')

    # Add a closing note paragraph
    doc.add_paragraph(
        'Please ensure each step is completed and verified before proceeding to the next. '
        'Contact the data engineering team if any issues arise during the migration.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
