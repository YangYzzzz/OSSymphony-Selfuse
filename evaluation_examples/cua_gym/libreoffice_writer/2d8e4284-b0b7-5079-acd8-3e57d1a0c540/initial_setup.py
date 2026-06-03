"""
Initial Setup: School Newsletter - A4 portrait, single column, default margins
Task ID: writer_page_070
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'school_newsletter'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup: A4, portrait, single column, margins 2.54cm all sides ---
    section = doc.sections[0]
    # A4 dimensions in portrait
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Single column (default - no multi-column XML needed)

    # --- Page 1: School Newsletter Content ---
    # Newsletter title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('WESTFIELD ELEMENTARY SCHOOL NEWSLETTER')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = subtitle.add_run('Volume 12 | Issue 3 | March 2025')
    run2.italic = True
    run2.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Principal's Message
    heading1 = doc.add_paragraph()
    run3 = heading1.add_run("Principal's Message")
    run3.bold = True
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph(
        "Dear Westfield Families, as we move through the spring semester, I am delighted "
        "to share the many exciting achievements and upcoming events at our school. Our "
        "students continue to inspire us every day with their enthusiasm for learning and "
        "their commitment to our school values of respect, responsibility, and resilience."
    )

    doc.add_paragraph(
        "This month we celebrated Read Across America Week with a series of special "
        "reading activities, author visits, and storytelling sessions. Over 320 students "
        "participated in our school-wide reading challenge, collectively logging more than "
        "4,500 hours of reading. Congratulations to Mrs. Nakamura's 4th grade class for "
        "the most hours logged in any single classroom!"
    )

    # Academic News
    heading2 = doc.add_paragraph()
    run4 = heading2.add_run('Academic News')
    run4.bold = True
    run4.font.size = Pt(14)
    run4.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph(
        "State Assessment Preparation: The spring state assessments are approaching. "
        "Students in grades 3-5 will take the English Language Arts and Mathematics "
        "assessments from April 14-18. Teachers have been working diligently to prepare "
        "students with targeted review sessions and practice materials."
    )

    doc.add_paragraph(
        "Science Fair Winners: Congratulations to all participants in this year's Science "
        "Fair! First place went to Amara Okafor (5th grade) for her project on water "
        "filtration systems. Second place was awarded to twins Diego and Sofia Reyes "
        "(4th grade) for their investigation into solar energy. Third place honors went "
        "to Liam Patel (3rd grade) for his study of local bird migration patterns."
    )

    # Sports & Activities
    heading3 = doc.add_paragraph()
    run5 = heading3.add_run('Sports & Extracurricular Activities')
    run5.bold = True
    run5.font.size = Pt(14)
    run5.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph(
        "Spring Sports: Our soccer and softball teams have kicked off their spring "
        "seasons with impressive performances. The Westfield Wolves soccer team is "
        "currently undefeated with a record of 4-0, led by coach Mr. Hernandez. "
        "The softball team opened with a strong 8-3 victory against Riverside Elementary."
    )

    doc.add_paragraph(
        "Drama Club Production: Mark your calendars! The Drama Club will present "
        "'The Wizard of Oz' on Friday, March 21st and Saturday, March 22nd at 6:30 PM "
        "in the school auditorium. Tickets are $5 for adults and $3 for students. "
        "All proceeds benefit the school arts program."
    )

    # Page break to page 2
    doc.add_page_break()

    # --- Page 2: More newsletter content ---
    heading4 = doc.add_paragraph()
    run6 = heading4.add_run('Community & Family Engagement')
    run6.bold = True
    run6.font.size = Pt(14)
    run6.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph(
        "Volunteer Appreciation Week: April 7-11 is our annual Volunteer Appreciation "
        "Week. This year we have over 85 active parent and community volunteers who "
        "contribute countless hours to support our students and staff. Please join us "
        "for a special recognition breakfast on Thursday, April 10th at 8:00 AM in "
        "the school cafeteria."
    )

    doc.add_paragraph(
        "Spring Book Fair: The Scholastic Book Fair returns to Westfield from March "
        "24-28. Students will have the opportunity to browse and purchase books during "
        "their scheduled library periods. Family night shopping will be held on "
        "Wednesday, March 26th from 5:00-7:30 PM. Wish lists can be set up online "
        "at scholastic.com/bookfairs."
    )

    # Upcoming Events
    heading5 = doc.add_paragraph()
    run7 = heading5.add_run('Upcoming Events Calendar')
    run7.bold = True
    run7.font.size = Pt(14)
    run7.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    events = [
        ('March 17', 'St. Patrick\'s Day - Spirit Day (wear green!)'),
        ('March 20', 'Spring Begins - Garden Club Planting Day'),
        ('March 21-22', 'Drama Club: The Wizard of Oz - 6:30 PM'),
        ('March 24-28', 'Scholastic Spring Book Fair'),
        ('March 26', 'Book Fair Family Night 5:00-7:30 PM'),
        ('April 4', 'End of 3rd Quarter'),
        ('April 7', 'Spring Break Begins'),
        ('April 14', 'Students Return from Spring Break'),
        ('April 14-18', 'State Assessments (Grades 3-5)'),
        ('April 25', 'Earth Day Environmental Fair'),
    ]

    for date, event in events:
        para = doc.add_paragraph()
        run_date = para.add_run(f'{date}: ')
        run_date.bold = True
        para.add_run(event)

    # Health & Wellness
    heading6 = doc.add_paragraph()
    run8 = heading6.add_run('Health & Wellness Reminder')
    run8.bold = True
    run8.font.size = Pt(14)
    run8.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph(
        "As allergy season approaches, please ensure the school nurse has up-to-date "
        "information about your child's allergies and any required medications. The "
        "nurse's office is open daily from 7:45 AM to 3:30 PM. Contact Nurse Patricia "
        "Williams at ext. 215 or pwilliams@westfield.edu with any health-related concerns."
    )

    # Contact Info
    heading7 = doc.add_paragraph()
    run9 = heading7.add_run('Contact Information')
    run9.bold = True
    run9.font.size = Pt(14)
    run9.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph(
        "Main Office: (555) 234-7890 | Principal Dr. Susan Hartley: ext. 101\n"
        "Attendance Line: (555) 234-7891 | Website: www.westfieldelementary.edu\n"
        "Newsletter Editor: Ms. Caroline Brooks, cbrooks@westfield.edu"
    )

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
