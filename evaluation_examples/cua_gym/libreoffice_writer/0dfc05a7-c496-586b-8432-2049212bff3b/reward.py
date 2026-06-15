"""
Reward Script: Create nested tables for quarterly data comparison
Task ID: writer_tm_045
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Outer 2x2 table exists in the document
  Component 2 (0.4): Each of the 4 outer cells contains a 3x3 nested table
  Component 3 (0.3): Nested tables contain non-empty data (header row + data rows)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_045'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Outer 2x2 table exists (0.3 points)
    # Initial env has 0 tables, golden env has 1 outer table with 2 rows x 2 cols
    try:
        if len(doc.tables) >= 1:
            outer_table = doc.tables[0]
            num_rows = len(outer_table.rows)
            num_cols = len(outer_table.columns)
            if num_rows == 2 and num_cols == 2:
                print(f"PASS: Component 1 — Outer 2x2 table found ({num_rows}x{num_cols}) (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — Outer table is {num_rows}x{num_cols}, expected 2x2")
        else:
            print(f"FAIL: Component 1 — No tables found in document (found {len(doc.tables)})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Each of the 4 outer cells contains a nested 3x3 table (0.4 points)
    # Award 0.1 per cell that has a correctly-sized nested table
    try:
        if len(doc.tables) >= 1:
            outer_table = doc.tables[0]
            nested_count = 0
            correct_size_count = 0
            for ri in range(min(len(outer_table.rows), 2)):
                for ci in range(min(len(outer_table.columns), 2)):
                    cell = outer_table.cell(ri, ci)
                    num_nested = len(cell.tables)
                    if num_nested >= 1:
                        nested_count += 1
                        nt = cell.tables[0]
                        nt_rows = len(nt.rows)
                        nt_cols = len(nt.columns)
                        if nt_rows == 3 and nt_cols == 3:
                            correct_size_count += 1
                            print(f"  Cell({ri},{ci}): nested 3x3 table found")
                        else:
                            print(f"  Cell({ri},{ci}): nested table is {nt_rows}x{nt_cols}, expected 3x3")
                    else:
                        print(f"  Cell({ri},{ci}): no nested table found")

            if correct_size_count == 4:
                print(f"PASS: Component 2 — All 4 cells have 3x3 nested tables (0.4 pts)")
                total_score += 0.4
            elif correct_size_count > 0:
                partial = round(0.1 * correct_size_count, 2)
                print(f"PARTIAL: Component 2 — {correct_size_count}/4 cells have 3x3 nested tables ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — No cells have correctly-sized nested tables (nested found: {nested_count})")
        else:
            print(f"FAIL: Component 2 — No outer table exists")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Nested tables contain non-empty data (0.3 points)
    # Each nested table should have content in its cells (not all empty)
    # Award 0.075 per nested table that has at least 4 non-empty cells (header + some data)
    try:
        if len(doc.tables) >= 1:
            outer_table = doc.tables[0]
            data_count = 0
            for ri in range(min(len(outer_table.rows), 2)):
                for ci in range(min(len(outer_table.columns), 2)):
                    cell = outer_table.cell(ri, ci)
                    if len(cell.tables) >= 1:
                        nt = cell.tables[0]
                        non_empty = 0
                        for nri in range(len(nt.rows)):
                            for nci in range(len(nt.columns)):
                                if nt.cell(nri, nci).text.strip():
                                    non_empty += 1
                        if non_empty >= 4:
                            data_count += 1
                            print(f"  Cell({ri},{ci}): nested table has {non_empty} non-empty cells")
                        else:
                            print(f"  Cell({ri},{ci}): nested table has only {non_empty} non-empty cells (need >=4)")

            if data_count == 4:
                print(f"PASS: Component 3 — All 4 nested tables contain data (0.3 pts)")
                total_score += 0.3
            elif data_count > 0:
                partial = round(0.075 * data_count, 3)
                print(f"PARTIAL: Component 3 — {data_count}/4 nested tables have data ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — No nested tables contain sufficient data")
        else:
            print(f"FAIL: Component 3 — No outer table exists")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state("libreoffice_writer")

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
