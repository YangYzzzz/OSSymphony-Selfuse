"""
Initial Setup: Create a document with heading for quarterly comparison (no tables yet).
Task ID: writer_tm_045
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add heading
    doc.add_heading('Quarterly Comparison', level=1)

    # Add some introductory context so the document is not empty
    para = doc.add_paragraph(
        'This document is intended for comparing quarterly performance data '
        'across four business divisions. Please create a structured table layout '
        'below to organize the quarterly figures.'
    )
    para.paragraph_format.space_after = Pt(12)

    para2 = doc.add_paragraph(
        'Each division should have its own mini-table with rows for key metrics '
        '(Revenue, Expenses, Net Profit) across three time periods.'
    )
    para2.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
