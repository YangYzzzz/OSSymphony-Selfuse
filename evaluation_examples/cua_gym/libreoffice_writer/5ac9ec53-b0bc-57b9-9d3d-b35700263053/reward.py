"""
FINAL REWARD SCRIPT - SUCCESS
Task: Put the "Summary" heading on a centered line.
Generated: 2025-10-14 11:40:48
Status: success
Model: azure-o3
Total Steps: 1
"""

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def verify_summary_centered(file_path: str) -> float:
    """Verify that the document contains a paragraph whose exact text is
    'Summary' (case-insensitive, trimmed) and that this paragraph is centered.

    Scoring (progressive):
        - 0.5 points  – A paragraph with text exactly 'Summary' is present.
        - 0.5 points  – That paragraph (at least one, if several) has CENTER alignment.
    Returns a float between 0.0 and 1.0.
    """

    print(f"Verifying file: {file_path}")

    # Safety check – file exists
    if not os.path.exists(file_path):
        print("✗ File does not exist")
        return 0.0  # No progress possible

    # Attempt to load the DOCX document
    try:
        doc = Document(file_path)
        print(f"✓ Document loaded with {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Error loading document: {e}")
        return 0.0

    score = 0.0

    # Locate paragraphs with text exactly 'Summary'
    summary_paras = [p for p in doc.paragraphs if p.text.strip().lower() == 'summary']
    if not summary_paras:
        print("✗ No 'Summary' paragraph found (0 points)")
        return score  # 0.0

    print(f"✓ Found {len(summary_paras)} 'Summary' paragraph(s)")
    score += 0.5  # Presence of required heading

    # Check alignment for at least one of those paragraphs
    centered_found = False
    for para in summary_paras:
        alignment = para.paragraph_format.alignment  # None = default/left
        align_label = 'None/Default' if alignment is None else WD_ALIGN_PARAGRAPH(alignment).name
        print(f"  Alignment: {align_label}")
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            centered_found = True
            break

    if centered_found:
        print("✓ 'Summary' paragraph is centered (0.5 points)")
        score += 0.5
    else:
        print("✗ 'Summary' paragraph is not centered (0 points)")

    final_score = min(score, 1.0)
    print(f"Final Score: {final_score}")
    return final_score


if __name__ == "__main__":
    # Path to the document provided in the task context
    file_path = "/home/user/put_the_summary_heading_on_a_centered_line.docx"
    reward = verify_summary_centered(file_path)
    print(f"REWARD: {reward}")
