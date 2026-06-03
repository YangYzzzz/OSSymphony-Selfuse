"""
Initial Setup: Research Summary document with four body paragraphs in 12pt Times New Roman.
Task ID: writer_txtfmt_067
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_067'
OUTPUT = f'{WORKDIR}/Desktop/research_summary.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Research Summary: Renewable Energy Adoption")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)

    # --- Paragraph 1 ---
    # First sentence will need to be bolded by agent
    p1 = doc.add_paragraph()
    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # agent must set JUSTIFY
    p1.paragraph_format.space_after = Pt(8)
    r1a = p1.add_run(
        "Solar energy adoption has increased dramatically over the past decade. "
    )
    r1a.font.name = "Times New Roman"
    r1a.font.size = Pt(12)
    r1b = p1.add_run(
        "Photovoltaic technology improvements have driven costs down by over 80% since 2010, "
        "making solar power cost-competitive with fossil fuels in many regions. "
        "Residential rooftop installations have surged across North America, Europe, and Asia-Pacific, "
        "with an estimated 100 gigawatts of new capacity added globally in 2023 alone."
    )
    r1b.font.name = "Times New Roman"
    r1b.font.size = Pt(12)

    # --- Paragraph 2 ---
    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p2.paragraph_format.space_after = Pt(8)
    r2a = p2.add_run(
        "Wind power generation now accounts for 12% of global electricity. "
    )
    r2a.font.name = "Times New Roman"
    r2a.font.size = Pt(12)
    r2b = p2.add_run(
        "Offshore wind farms have expanded significantly, with turbines now capable of "
        "generating up to 15 megawatts each. "
        "Countries such as Denmark, the United Kingdom, and Germany have made substantial "
        "investments in offshore wind infrastructure to meet their renewable energy targets. "
        "Supply chain improvements and larger turbine designs continue to reduce the levelized "
        "cost of energy for wind power."
    )
    r2b.font.name = "Times New Roman"
    r2b.font.size = Pt(12)

    # --- Paragraph 3 ---
    p3 = doc.add_paragraph()
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p3.paragraph_format.space_after = Pt(8)
    r3a = p3.add_run(
        "Battery storage technology remains the primary bottleneck. "
    )
    r3a.font.name = "Times New Roman"
    r3a.font.size = Pt(12)
    r3b = p3.add_run(
        "While lithium-ion battery costs have fallen dramatically, grid-scale energy storage "
        "still struggles to keep pace with the intermittent nature of solar and wind generation. "
        "Research into solid-state batteries, flow batteries, and other next-generation storage "
        "solutions is accelerating, driven by increasing demand from utilities and electric vehicle "
        "manufacturers who require higher energy density, faster charging, and longer cycle lives."
    )
    r3b.font.name = "Times New Roman"
    r3b.font.size = Pt(12)

    # --- Paragraph 4 ---
    p4 = doc.add_paragraph()
    p4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p4.paragraph_format.space_after = Pt(8)
    r4a = p4.add_run(
        "Government subsidies play a crucial role in renewable energy adoption. "
    )
    r4a.font.name = "Times New Roman"
    r4a.font.size = Pt(12)
    r4b = p4.add_run(
        "Tax incentives, feed-in tariffs, and direct grants have accelerated deployment "
        "in markets that would otherwise have taken years longer to reach economic viability. "
        "Policy frameworks like the US Inflation Reduction Act and the EU Green Deal have "
        "committed hundreds of billions of dollars to clean energy transition programs, "
        "stimulating private investment and creating millions of jobs in the renewable sector."
    )
    r4b.font.name = "Times New Roman"
    r4b.font.size = Pt(12)

    # Ensure Desktop directory exists on VM (handled by script on VM)
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
