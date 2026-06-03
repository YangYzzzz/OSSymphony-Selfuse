"""
Reward Script: Make all paragraph text justified and bold first sentence of each paragraph
Task ID: writer_txtfmt_067
Domain: libreoffice_writer
Scoring:
  Component 1: All 4 body paragraphs have JUSTIFY alignment (0.4 pts)
  Component 2: First sentence (Run 0) of all 4 body paragraphs is bold=True (0.4 pts)
  Component 3: Remaining runs in all 4 body paragraphs are NOT bold — only awarded when
               Component 2 also passes (ensures complete task execution, not just partial) (0.2 pts)
Total: 1.0

NOTE: Component 3 is gated on Component 2 to avoid awarding points for the pre-existing
non-bold state in the initial artifact (the task only introduces bold on first sentences).
"""

import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'research_summary'

# The 4 body paragraphs are at indices 1, 2, 3, 4 (index 0 is the title)
BODY_PARA_INDICES = [1, 2, 3, 4]

# Expected first sentence starters — used to confirm we are checking the right paragraphs
FIRST_SENTENCE_STARTERS = [
    'Solar energy adoption',
    'Wind power generation',
    'Battery storage technology',
    'Government subsidies',
]


def count_justified_paragraphs(doc):
    """Count how many of the 4 body paragraphs have JUSTIFY alignment."""
    count = 0
    for idx in BODY_PARA_INDICES:
        para = doc.paragraphs[idx]
        if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            count += 1
        else:
            print(f"FAIL: Component 1 — Para {idx} alignment={para.paragraph_format.alignment}, expected JUSTIFY")
    return count


def count_bold_first_sentences(doc):
    """Count how many of the 4 body paragraphs have Run 0 explicitly bold=True."""
    count = 0
    for i, idx in enumerate(BODY_PARA_INDICES):
        para = doc.paragraphs[idx]
        if len(para.runs) == 0:
            print(f"FAIL: Component 2 — Para {idx} has no runs")
            continue
        run0 = para.runs[0]
        expected_starter = FIRST_SENTENCE_STARTERS[i]
        if not run0.text.strip().startswith(expected_starter[:15]):
            print(f"FAIL: Component 2 — Para {idx} Run 0 unexpected text: {repr(run0.text[:50])}")
            continue
        if run0.bold is True:
            count += 1
        else:
            print(f"FAIL: Component 2 — Para {idx} Run 0 bold={run0.bold}, expected True")
    return count


def count_paragraphs_with_correct_non_bold_rest(doc):
    """Count how many of the 4 body paragraphs have non-bold (not explicitly True) Run 1+."""
    count = 0
    for idx in BODY_PARA_INDICES:
        para = doc.paragraphs[idx]
        runs_after_first = para.runs[1:]
        if not runs_after_first:
            count += 1
            continue
        all_non_bold = all(
            run.bold is not True
            for run in runs_after_first
            if run.text.strip()
        )
        if all_non_bold:
            count += 1
        else:
            print(f"FAIL: Component 3 — Para {idx} has a non-first run that is bold")
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify the document has at least 5 paragraphs
    if len(doc.paragraphs) < 5:
        print(f"CRITICAL: Expected at least 5 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 4 body paragraphs have JUSTIFY alignment (0.4 points)
    # In initial_env all body paragraphs are LEFT aligned — FAILS on initial, PASSES on golden
    try:
        justified_count = count_justified_paragraphs(doc)
        if justified_count == 4:
            print(f"PASS: Component 1 — All 4 body paragraphs are JUSTIFY aligned (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Only {justified_count}/4 body paragraphs are JUSTIFY aligned")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: First sentence run (Run 0) of all 4 body paragraphs is bold=True (0.4 points)
    # In initial_env Run 0 bold=None (not bold) — FAILS on initial, PASSES on golden
    try:
        bold_first_count = count_bold_first_sentences(doc)
        if bold_first_count == 4:
            print(f"PASS: Component 2 — First sentence of all 4 body paragraphs is bold (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Only {bold_first_count}/4 paragraphs have bold first sentence")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")
        bold_first_count = 0

    # Component 3: Remaining text (Run 1+) of all 4 body paragraphs is NOT bold (0.2 points)
    # GATED on Component 2 passing: only evaluated when all 4 first sentences are correctly bold.
    # This prevents awarding points for pre-existing non-bold remaining-sentence state
    # in the initial artifact (where bold was never applied at all).
    if bold_first_count == 4:
        try:
            non_bold_rest_count = count_paragraphs_with_correct_non_bold_rest(doc)
            if non_bold_rest_count == 4:
                print(f"PASS: Component 3 — Remaining sentences in all 4 paragraphs are NOT bold (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Only {non_bold_rest_count}/4 paragraphs have correct non-bold remaining text")
        except Exception as e:
            print(f"ERROR: Component 3 — {e}")
    else:
        print("SKIP: Component 3 — skipped because Component 2 did not pass (prevents false positives on initial_env)")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
