"""
Initial Setup: Read editing notes from my_notes.docx and apply edits to screenshot.png using GIMP
Task ID: osworld_multi_apps_writer_to_gimp_011
Domain: multi_apps (libreoffice_writer + gimp)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_to_gimp_011'

NOTES_FILE = f'{DESKTOP}/my_notes.docx'
SCREENSHOT_FILE = f'{DESKTOP}/screenshot.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_notes_docx():
    """Create my_notes.docx with image editing instructions."""
    doc = Document()

    # Title
    title = doc.add_heading('Image Editing Notes', level=1)

    # Introduction paragraph
    doc.add_paragraph(
        'These are my personal notes for editing screenshot.png. '
        'Please follow the steps below carefully when editing the image using GIMP.'
    )

    # Step 1: Crop
    doc.add_heading('Step 1: Crop', level=2)
    p1 = doc.add_paragraph()
    run1 = p1.add_run('Crop the top 100 pixels from the image.')
    run1.font.size = Pt(12)

    doc.add_paragraph(
        'The top portion of the screenshot contains the system status bar which '
        'is not relevant to the content. Remove it by cropping 100 pixels from the top.'
    )

    # Step 2: Blue Overlay
    doc.add_heading('Step 2: Add Blue Overlay', level=2)
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Add a 50% opacity blue overlay to the entire image.')
    run2.font.size = Pt(12)

    doc.add_paragraph(
        'Apply a semi-transparent blue color layer (RGB: 0, 0, 255) at 50% opacity '
        'over the entire cropped image. This creates a blue-tinted effect while '
        'keeping the underlying image visible.'
    )

    # Step 3: Save
    doc.add_heading('Step 3: Save Result', level=2)
    p3 = doc.add_paragraph()
    run3 = p3.add_run('Save the edited image as screenshot_edited.png on the Desktop.')
    run3.font.size = Pt(12)

    doc.add_paragraph(
        'Export/save the final result as: /home/user/Desktop/screenshot_edited.png\n'
        'Use PNG format to preserve quality.'
    )

    # Summary note
    doc.add_heading('Summary', level=2)
    summary = doc.add_paragraph()
    summary_run = summary.add_run(
        'In short:\n'
        '1. Crop top 100 pixels\n'
        '2. Add 50% opacity blue (0, 0, 255) overlay\n'
        '3. Save as screenshot_edited.png'
    )
    summary_run.font.bold = True

    os.makedirs(DESKTOP, exist_ok=True)
    doc.save(NOTES_FILE)
    print(f'Created notes file: {NOTES_FILE}')


def create_screenshot_png():
    """Create a realistic-looking screenshot.png on the Desktop."""
    # Create a simulated desktop screenshot: 1280x800
    width, height = 1280, 800
    img = Image.new('RGB', (width, height), color=(240, 240, 240))
    draw = ImageDraw.Draw(img)

    # Draw a simulated taskbar at top (status bar)
    draw.rectangle([0, 0, width, 30], fill=(45, 45, 48))
    # Time on right side of taskbar
    draw.rectangle([width - 120, 5, width - 10, 25], fill=(70, 70, 75))

    # Draw a window with title bar
    draw.rectangle([50, 60, 900, 600], fill=(255, 255, 255), outline=(180, 180, 180), width=2)
    # Window title bar
    draw.rectangle([50, 60, 900, 90], fill=(0, 120, 215))
    # Window close/min/max buttons
    draw.ellipse([870, 70, 885, 85], fill=(232, 17, 35))
    draw.ellipse([848, 70, 863, 85], fill=(255, 185, 0))
    draw.ellipse([826, 70, 841, 85], fill=(0, 200, 0))

    # Window content area - simulated document text (fixed values for reproducibility)
    line_data = [
        (70, 110, 70 + 700, 110 + 14, 80),
        (70, 138, 70 + 620, 138 + 14, 120),
        (70, 166, 70 + 450, 166 + 14, 120),
        (70, 194, 70 + 680, 194 + 14, 120),
        (70, 222, 70 + 300, 222 + 14, 120),
        (70, 250, 70 + 720, 250 + 14, 120),
        (70, 278, 70 + 390, 278 + 14, 120),
        (70, 306, 70 + 560, 306 + 14, 120),
        (70, 334, 70 + 640, 334 + 14, 120),
        (70, 362, 70 + 480, 362 + 14, 120),
        (70, 390, 70 + 710, 390 + 14, 120),
        (70, 418, 70 + 340, 418 + 14, 120),
        (70, 446, 70 + 590, 446 + 14, 120),
        (70, 474, 70 + 430, 474 + 14, 120),
        (70, 502, 70 + 660, 502 + 14, 120),
        (70, 530, 70 + 510, 530 + 14, 120),
    ]
    for x1, y1, x2, y2, gray_val in line_data:
        draw.rectangle([x1, y1, x2, y2], fill=(gray_val, gray_val, gray_val))

    # Draw a sidebar panel
    draw.rectangle([920, 60, 1200, 600], fill=(249, 249, 249), outline=(220, 220, 220), width=1)
    draw.rectangle([920, 60, 1200, 85], fill=(230, 230, 230))
    for i, y_offset in enumerate(range(100, 580, 40)):
        draw.rectangle([935, y_offset, 1185, y_offset + 25], fill=(210, 210, 210))

    # Desktop icons at bottom
    for i in range(5):
        x = 60 + i * 80
        draw.rectangle([x, 650, x + 50, 720], fill=(200, 200, 255), outline=(150, 150, 200))

    # Taskbar at bottom
    draw.rectangle([0, 755, width, 800], fill=(30, 30, 30))
    draw.rectangle([10, 760, 60, 795], fill=(0, 120, 215))  # Start button
    for i in range(6):
        x = 80 + i * 50
        draw.rectangle([x, 760, x + 40, 795], fill=(60, 60, 60), outline=(80, 80, 80))

    os.makedirs(DESKTOP, exist_ok=True)
    img.save(SCREENSHOT_FILE, 'PNG')
    print(f'Created screenshot: {SCREENSHOT_FILE}')


def create_initial():
    # Create the notes document
    create_notes_docx()

    # Create the screenshot PNG
    create_screenshot_png()

    # GUI startup: open my_notes.docx in LibreOffice Writer for agent to read
    launch_gui(f'libreoffice --writer "{NOTES_FILE}"', delay_sec=2.0)
    # Also give agent access to GIMP for editing
    # (Agent will open GIMP itself after reading the notes)
    print('GUI_READY: launched LibreOffice Writer with my_notes.docx using DISPLAY=:0')


create_initial()
