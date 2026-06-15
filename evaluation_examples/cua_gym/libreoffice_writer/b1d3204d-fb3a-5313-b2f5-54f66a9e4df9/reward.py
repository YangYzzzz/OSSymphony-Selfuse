"""
Reward Script: Business card layout using Avery 8371 format
Task ID: writer_lec_045
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Table exists with 5 rows x 2 columns (10 cards)
  Component 2 (0.30): All 10 cells contain required contact info
  Component 3 (0.20): Name "Robert Chen" is bold in cards
  Component 4 (0.20): Card dimensions approximately 3.5 x 2 inches
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_045'

# Required contact info fields (case-insensitive substring matching)
REQUIRED_FIELDS = [
    'robert chen',
    'senior developer',
    'devops solutions',
    '(555) 123-4567',
    'robert@devops.com',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table with 5 rows x 2 columns (0.30 points)
    # Initial file has 0 tables, golden has a 5x2 table
    try:
        if len(doc.tables) < 1:
            print(f"FAIL: Component 1 — No tables found (found {len(doc.tables)})")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 5 and num_cols == 2:
                print(f"PASS: Component 1 — Table has 5 rows x 2 columns (10 cards) (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 1 — Table is {num_rows}x{num_cols}, expected 5x2")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 10 cells contain required contact info (0.30 points)
    # Each cell should have all 5 fields. Award partial: 0.03 per correct cell.
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 2 — No tables to check")
        else:
            table = doc.tables[0]
            cells_correct = 0
            total_cells = 0
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    total_cells += 1
                    cell_text = cell.text.lower()
                    all_present = all(field in cell_text for field in REQUIRED_FIELDS)
                    if all_present:
                        cells_correct += 1
                    else:
                        missing = [f for f in REQUIRED_FIELDS if f not in cell_text]
                        print(f"  Cell({ri},{ci}) missing: {missing}")

            if total_cells > 0 and cells_correct == 10:
                print(f"PASS: Component 2 — All 10 cells contain full contact info (0.30 pts)")
                total_score += 0.30
            elif cells_correct > 0:
                partial = round(0.03 * cells_correct, 2)
                print(f"PARTIAL: Component 2 — {cells_correct}/{total_cells} cells correct ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — 0 cells contain required contact info")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Name "Robert Chen" is bold in cards (0.20 points)
    # Check that in each cell, the run containing "Robert Chen" has bold=True
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 3 — No tables to check")
        else:
            table = doc.tables[0]
            bold_count = 0
            checked_cells = 0
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if 'robert chen' not in cell.text.lower():
                        continue
                    checked_cells += 1
                    name_bold = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if 'robert chen' in run.text.lower().strip():
                                if run.font.bold is True:
                                    name_bold = True
                                    break
                        if name_bold:
                            break
                    if name_bold:
                        bold_count += 1

            if checked_cells >= 10 and bold_count >= 10:
                print(f"PASS: Component 3 — 'Robert Chen' is bold in all 10 cards (0.20 pts)")
                total_score += 0.20
            elif bold_count > 0:
                partial = round(0.02 * bold_count, 2)
                print(f"PARTIAL: Component 3 — Bold name in {bold_count}/{checked_cells} cells ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Name not bold in any cell (checked {checked_cells} cells)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Card dimensions approximately 3.5 x 2 inches (0.20 points)
    # Avery 8371: 3.5" wide x 2" tall. Allow tolerance of +/- 0.5 inches.
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 4 — No tables to check")
        else:
            table = doc.tables[0]
            EMU_PER_INCH = 914400
            # Check first cell width and first row height
            cell_width = table.rows[0].cells[0].width
            row_height = table.rows[0].height

            width_ok = False
            height_ok = False

            if cell_width is not None:
                width_in = cell_width / EMU_PER_INCH
                if 3.0 <= width_in <= 4.0:
                    width_ok = True
                    print(f"  Width: {width_in:.2f}in (target ~3.5in) — OK")
                else:
                    print(f"  Width: {width_in:.2f}in (target ~3.5in) — out of range")
            else:
                print("  Width: None (cannot verify)")

            if row_height is not None:
                height_in = row_height / EMU_PER_INCH
                if 1.5 <= height_in <= 2.5:
                    height_ok = True
                    print(f"  Height: {height_in:.2f}in (target ~2.0in) — OK")
                else:
                    print(f"  Height: {height_in:.2f}in (target ~2.0in) — out of range")
            else:
                print("  Height: None (cannot verify)")

            if width_ok and height_ok:
                print(f"PASS: Component 4 — Card dimensions match Avery 8371 (0.20 pts)")
                total_score += 0.20
            elif width_ok or height_ok:
                print(f"PARTIAL: Component 4 — One dimension matches (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Dimensions do not match Avery 8371")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
