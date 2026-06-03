"""
Reward Script: Laptop Comparison Guide 2025 — product comparison document
Task ID: writer_wf_033
Domain: libreoffice_writer
Scoring:
  C1 (0.15): Title "Laptop Comparison Guide 2025"
  C2 (0.20): Table with 11 rows x 4 columns
  C3 (0.15): Header row content (Feature, Model A/B/C)
  C4 (0.10): 10 feature row labels present
  C5 (0.15): Header row bold + white text
  C6 (0.10): Header row dark blue shading
  C7 (0.15): Three recommendation paragraphs
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_033'


def color_distance_hex(hex1, hex2):
    """Euclidean distance between two hex color strings (e.g. '003366')."""
    r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
    r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Missing dependency: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title "Laptop Comparison Guide 2025" (0.15 points)
    try:
        title_matches = [para for para in doc.paragraphs
                         if 'Laptop Comparison Guide 2025' in para.text]
        if len(title_matches) > 0:
            print(f"PASS: Component 1 — Title 'Laptop Comparison Guide 2025' found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title 'Laptop Comparison Guide 2025' not found in any paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table with 11 rows x 4 columns (0.20 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            nrows = len(table.rows)
            ncols = len(table.columns)
            if nrows == 11 and ncols == 4:
                print(f"PASS: Component 2 — Table has 11 rows x 4 columns (0.20 pts)")
                total_score += 0.20
            elif nrows >= 10 and ncols == 4:
                # Partial credit: close but not exact row count
                print(f"PARTIAL: Component 2 — Table has {nrows} rows x {ncols} cols (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 2 — Table has {nrows} rows x {ncols} cols, expected 11x4")
        else:
            print(f"FAIL: Component 2 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Gate: need at least one table to proceed with table checks
    if len(doc.tables) == 0:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    table = doc.tables[0]

    # Component 3: Header row content (0.15 points)
    try:
        expected_headers = ['Feature', 'Model A (ProBook X1)', 'Model B (TechNote 15)', 'Model C (UltraSlim Z)']
        actual_headers = [table.rows[0].cells[ci].text.strip() for ci in range(min(len(table.columns), 4))]
        matches = sum(1 for exp, act in zip(expected_headers, actual_headers) if exp.lower() in act.lower())
        if matches == 4:
            print(f"PASS: Component 3 — All 4 header cells match (0.15 pts)")
            total_score += 0.15
        elif matches >= 2:
            pts = round(0.15 * matches / 4, 2)
            print(f"PARTIAL: Component 3 — {matches}/4 header cells match ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 3 — Headers: {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 10 feature row labels (0.10 points)
    try:
        expected_features = ['Processor', 'RAM', 'Storage', 'Display', 'Battery',
                             'Weight', 'Price', 'Warranty', 'Rating', 'Best For']
        actual_features = []
        for ri in range(1, min(len(table.rows), 12)):
            actual_features.append(table.rows[ri].cells[0].text.strip())

        found = sum(1 for ef in expected_features
                    if any(ef.lower() in af.lower() for af in actual_features))
        if found >= 9:
            print(f"PASS: Component 4 — {found}/10 feature labels found (0.10 pts)")
            total_score += 0.10
        elif found >= 5:
            pts = round(0.10 * found / 10, 2)
            print(f"PARTIAL: Component 4 — {found}/10 feature labels ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 4 — Only {found}/10 feature labels found. Actual: {actual_features}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Header row bold + white text (0.15 points)
    try:
        header_bold_count = 0
        header_white_count = 0
        total_header_runs = 0
        for ci in range(min(len(table.columns), 4)):
            cell = table.rows[0].cells[ci]
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        total_header_runs += 1
                        if run.font.bold:
                            header_bold_count += 1
                        # Check white text color
                        if run.font.color and run.font.color.rgb:
                            rgb_str = str(run.font.color.rgb).upper()
                            if rgb_str == 'FFFFFF' or color_distance_hex(rgb_str, 'FFFFFF') < 30:
                                header_white_count += 1

        if total_header_runs > 0:
            bold_ok = header_bold_count == total_header_runs
            white_ok = header_white_count == total_header_runs
            if bold_ok and white_ok:
                print(f"PASS: Component 5 — Header row: all {total_header_runs} runs bold + white (0.15 pts)")
                total_score += 0.15
            elif bold_ok or white_ok:
                print(f"PARTIAL: Component 5 — bold={bold_ok}, white={white_ok} (0.075 pts)")
                total_score += 0.075
            else:
                print(f"FAIL: Component 5 — bold={header_bold_count}/{total_header_runs}, white={header_white_count}/{total_header_runs}")
        else:
            print(f"FAIL: Component 5 — No text runs in header row")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Header row dark blue shading (0.10 points)
    try:
        shaded_count = 0
        for ci in range(min(len(table.columns), 4)):
            tc = table.rows[0].cells[ci]._tc
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is not None:
                shd = tc_pr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill:
                        fill = fill.upper().lstrip('#')
                        # Dark blue: check if the color is in the blue family and dark
                        # Parse RGB
                        r = int(fill[0:2], 16)
                        g = int(fill[2:4], 16)
                        b = int(fill[4:6], 16)
                        # Dark blue: blue dominant, low red/green, blue > 50
                        if b > r and b > g and b >= 50 and r < 150 and g < 150:
                            shaded_count += 1
                        else:
                            print(f"  Cell[0,{ci}]: fill={fill}, not dark blue (r={r},g={g},b={b})")
        if shaded_count == 4:
            print(f"PASS: Component 6 — All 4 header cells have dark blue shading (0.10 pts)")
            total_score += 0.10
        elif shaded_count >= 2:
            pts = round(0.10 * shaded_count / 4, 2)
            print(f"PARTIAL: Component 6 — {shaded_count}/4 cells dark blue ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 6 — Only {shaded_count}/4 cells have dark blue shading")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Three recommendation paragraphs (0.15 points)
    try:
        model_names = ['ProBook X1', 'TechNote 15', 'UltraSlim Z']
        recs_found = 0
        for model in model_names:
            for para in doc.paragraphs:
                text = para.text.strip()
                if model.lower() in text.lower() and 'recommend' in text.lower():
                    recs_found += 1
                    break

        if recs_found == 3:
            print(f"PASS: Component 7 — All 3 recommendation paragraphs found (0.15 pts)")
            total_score += 0.15
        elif recs_found >= 1:
            pts = round(0.15 * recs_found / 3, 2)
            print(f"PARTIAL: Component 7 — {recs_found}/3 recommendation paragraphs ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 7 — No recommendation paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

import time
time.sleep(0.5)

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
