"""
Reward Script: Competitive Feature Matrix Table Creation
Task ID: writer_mktg_043
Domain: libreoffice_writer
Scoring:
  - Component 1: Table exists with correct dimensions (9x5)            0.20
  - Component 2: Header row: correct names, bold, dark bg, white text  0.25
  - Component 3: Nexus Platform column highlighted (#E8F5E9)           0.15
  - Component 4: All 8 features correctly populated (checkmarks/X)     0.25
  - Component 5: Data cells color-coded (green/red) + centered         0.15
  Total: 1.0
"""

import os
from math import sqrt

FILE_PATH = '/home/user/Desktop/competitive_positioning.docx'
TASK_ID = 'writer_mktg_043'

# Color constants (hex string as used in docx shading fill attribute)
HEADER_BG_HEX = '263238'       # dark teal/charcoal
NEXUS_BG_HEX = 'E8F5E9'        # light green

# Expected positive/negative patterns per feature (for columns: Nexus, AlphaForce, BetaHub, GammaWorks)
EXPECTED_DATA = {
    'Real-time Analytics':              [True, True, False, True],
    'Custom Dashboards':                [True, False, True, False],
    'API Access':                       [True, True, True, False],
    'Mobile App Support':               [True, True, False, False],
    'SSO / SAML Integration':           [True, False, True, True],
    'Role-Based Access Control':        [True, True, False, False],
    'Data Export (CSV/Excel)':          [True, False, True, True],
    'Dedicated Customer Success Manager': [True, False, False, True],
}

EXPECTED_HEADER = ['Feature', 'Nexus Platform', 'AlphaForce', 'BetaHub', 'GammaWorks']

POSITIVE_MARKERS = {'\u2713', '\u2714', '\u2611', 'yes', 'Yes', 'YES', 'checkmark'}
NEGATIVE_MARKERS = {'X', 'x', 'No', 'no', 'NO', '\u2717', '\u2718'}


def rgb_distance(hex1, hex2):
    """Euclidean distance between two hex RGB colors."""
    r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
    r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def get_cell_fill(cell):
    """Return cell background fill hex string or None."""
    from docx.oxml.ns import qn
    tc = cell._tc
    shd = tc.find('.//' + qn('w:shd'))
    if shd is not None:
        fill = shd.get(qn('w:fill'))
        if fill and fill.upper() not in ('AUTO', 'FFFFFF', 'NONE', ''):
            return fill.upper()
    return None


def is_positive_marker(text):
    return text.strip() in POSITIVE_MARKERS


def is_negative_marker(text):
    return text.strip() in NEGATIVE_MARKERS


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: file must be a docx with content
    if not doc.paragraphs:
        print("CRITICAL: Document is empty")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------
    # Component 1: Table exists with 9 rows x 5 columns (0.20 points)
    # This fails on initial_env (0 tables) and passes on golden_env (1 table)
    # -------------------------------------------------------------------
    try:
        tables = doc.tables
        table_found = False
        target_table = None

        for t in tables:
            if len(t.rows) == 9 and len(t.columns) == 5:
                table_found = True
                target_table = t
                break

        if table_found:
            print("PASS: Component 1 -- Table with 9 rows x 5 columns found (0.20 pts)")
            total_score += 0.20
        else:
            row_counts = [len(t.rows) for t in tables] if tables else []
            col_counts = [len(t.columns) for t in tables] if tables else []
            print(f"FAIL: Component 1 -- No 9x5 table found. Tables: {len(tables)}, "
                  f"row counts: {row_counts}, col counts: {col_counts}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    if target_table is None:
        # Cannot continue without a table
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # -------------------------------------------------------------------
    # Component 2: Header row has correct product names, bold text,
    #              dark background (~#263238), white text (~#FFFFFF)  (0.25 points)
    # -------------------------------------------------------------------
    try:
        header_row = target_table.rows[0]
        header_names_ok = True
        header_bold_ok = True
        header_bg_ok = True
        header_white_ok = True
        issues = []

        for col_idx, cell in enumerate(header_row.cells):
            cell_text = cell.text.strip()
            expected_name = EXPECTED_HEADER[col_idx]

            # Check name
            if cell_text != expected_name:
                header_names_ok = False
                issues.append(f"col {col_idx}: expected '{expected_name}', got '{cell_text}'")

            # Check background fill (should be close to #263238)
            fill = get_cell_fill(cell)
            if fill is None:
                header_bg_ok = False
                issues.append(f"col {col_idx}: missing dark background fill")
            elif rgb_distance(fill, HEADER_BG_HEX) > 50:
                header_bg_ok = False
                issues.append(f"col {col_idx}: bg fill '{fill}' too far from expected '{HEADER_BG_HEX}'")

            # Check run formatting: bold + white text
            para = cell.paragraphs[0]
            for run in para.runs:
                if not run.bold:
                    header_bold_ok = False
                    issues.append(f"col {col_idx}: run '{run.text}' not bold")
                color = run.font.color.rgb if run.font.color else None
                if color is None:
                    header_white_ok = False
                    issues.append(f"col {col_idx}: no explicit text color set")
                elif rgb_distance(str(color), 'FFFFFF') > 50:
                    header_white_ok = False
                    issues.append(f"col {col_idx}: text color '{color}' not white")

        if header_names_ok and header_bold_ok and header_bg_ok and header_white_ok:
            print("PASS: Component 2 -- Header row: correct names, bold, dark bg, white text (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 -- Header row issues: {'; '.join(issues[:5])}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # -------------------------------------------------------------------
    # Component 3: Nexus Platform column (col index 1) has light green
    #              background (#E8F5E9) in all 8 data rows  (0.15 points)
    # -------------------------------------------------------------------
    try:
        nexus_col_ok = True
        nexus_issues = []

        for row_idx in range(1, 9):  # rows 1..8
            cell = target_table.cell(row_idx, 1)
            fill = get_cell_fill(cell)
            if fill is None:
                nexus_col_ok = False
                nexus_issues.append(f"row {row_idx}: no fill")
            elif rgb_distance(fill, NEXUS_BG_HEX) > 50:
                nexus_col_ok = False
                nexus_issues.append(f"row {row_idx}: fill '{fill}' (expected ~'{NEXUS_BG_HEX}')")

        if nexus_col_ok:
            print("PASS: Component 3 -- Nexus Platform column has green background in all data rows (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 -- Nexus column bg issues: {'; '.join(nexus_issues[:3])}")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # -------------------------------------------------------------------
    # Component 4: All 8 features correctly populated with positive/negative
    #              markers matching the expected feature matrix  (0.25 points)
    # -------------------------------------------------------------------
    try:
        feature_issues = []
        matched_features = 0

        for row_idx in range(1, 9):
            row = target_table.rows[row_idx]
            feature_name = row.cells[0].text.strip()
            expected = EXPECTED_DATA.get(feature_name)
            if expected is None:
                feature_issues.append(f"row {row_idx}: unexpected feature '{feature_name}'")
                continue

            row_ok = True
            for col_idx in range(1, 5):
                cell_text = row.cells[col_idx].text.strip()
                expected_positive = expected[col_idx - 1]
                if expected_positive:
                    if not is_positive_marker(cell_text):
                        row_ok = False
                        feature_issues.append(
                            f"'{feature_name}' col {col_idx}: expected positive, got '{cell_text}'"
                        )
                else:
                    if not is_negative_marker(cell_text):
                        row_ok = False
                        feature_issues.append(
                            f"'{feature_name}' col {col_idx}: expected negative, got '{cell_text}'"
                        )
            if row_ok:
                matched_features += 1

        if matched_features == 8 and not feature_issues:
            print("PASS: Component 4 -- All 8 features correctly populated with checkmarks/X (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 -- {matched_features}/8 features correct. "
                  f"Issues: {'; '.join(feature_issues[:3])}")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # -------------------------------------------------------------------
    # Component 5: Data cells (rows 1-8, cols 1-4) are centered AND
    #              use green color for positive, red for negative  (0.15 points)
    # -------------------------------------------------------------------
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        color_align_issues = []
        cells_checked = 0
        cells_ok = 0

        # Expected: green ~#27AE60 for positive, red ~#C0392B for negative
        GREEN_HEX = '27AE60'
        RED_HEX = 'C0392B'

        for row_idx in range(1, 9):
            for col_idx in range(1, 5):
                cell = target_table.cell(row_idx, col_idx)
                para = cell.paragraphs[0]
                cells_checked += 1

                cell_ok = True

                # Check alignment (CENTER)
                align = para.paragraph_format.alignment
                if align != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    cell_ok = False
                    color_align_issues.append(
                        f"[{row_idx},{col_idx}]: alignment={align} (expected CENTER)"
                    )

                # Check text color
                for run in para.runs:
                    color = run.font.color.rgb if run.font.color else None
                    if color is None:
                        cell_ok = False
                        color_align_issues.append(f"[{row_idx},{col_idx}]: no text color")
                    else:
                        color_str = str(color).upper()
                        dist_green = rgb_distance(color_str, GREEN_HEX)
                        dist_red = rgb_distance(color_str, RED_HEX)
                        # Must be clearly green or red
                        if dist_green > 80 and dist_red > 80:
                            cell_ok = False
                            color_align_issues.append(
                                f"[{row_idx},{col_idx}]: color '{color_str}' is neither green nor red"
                            )

                if cell_ok:
                    cells_ok += 1

        if cells_ok == cells_checked and not color_align_issues:
            print("PASS: Component 5 -- All data cells centered with green/red color coding (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 -- {cells_ok}/{cells_checked} cells ok. "
                  f"Issues: {'; '.join(color_align_issues[:3])}")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
