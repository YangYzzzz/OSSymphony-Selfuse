"""
Initial Setup: Competitive Feature Matrix document (pre-table state)
Task ID: writer_mktg_043
Domain: libreoffice_writer

Creates competitive_positioning.docx with title, context paragraph, and
feature data as plain text. No table is created here — the agent must create it.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'competitive_positioning'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('Competitive Feature Matrix', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Context paragraph
    context_para = doc.add_paragraph(
        'The following analysis provides a direct comparison of our flagship product, '
        'Nexus Platform, against three leading competitors in the enterprise analytics market: '
        'AlphaForce, BetaHub, and GammaWorks. This matrix highlights key differentiators '
        'across eight critical feature categories that inform our Q3 go-to-market strategy.'
    )
    context_para.paragraph_format.space_after = Pt(12)

    # Subheading
    doc.add_heading('Feature Comparison Data', level=2)

    # Feature data as plain text (no table — agent must convert this)
    feature_lines = [
        'Real-time Analytics: Nexus (Yes), AlphaForce (Yes), BetaHub (No), GammaWorks (Yes).',
        'Custom Dashboards: Nexus (Yes), AlphaForce (No), BetaHub (Yes), GammaWorks (No).',
        'API Access: Nexus (Yes), AlphaForce (Yes), BetaHub (Yes), GammaWorks (No).',
        'Mobile App Support: Nexus (Yes), AlphaForce (Yes), BetaHub (No), GammaWorks (No).',
        'SSO / SAML Integration: Nexus (Yes), AlphaForce (No), BetaHub (Yes), GammaWorks (Yes).',
        'Role-Based Access Control: Nexus (Yes), AlphaForce (Yes), BetaHub (No), GammaWorks (No).',
        'Data Export (CSV/Excel): Nexus (Yes), AlphaForce (No), BetaHub (Yes), GammaWorks (Yes).',
        'Dedicated Customer Success Manager: Nexus (Yes), AlphaForce (No), BetaHub (No), GammaWorks (Yes).',
    ]

    for line in feature_lines:
        p = doc.add_paragraph(line, style='List Bullet')

    # Notes section
    doc.add_paragraph('')
    notes = doc.add_paragraph(
        'Note: This comparison is based on publicly available documentation and internal QA '
        'testing conducted in February 2025. Entries marked "Yes" indicate full feature '
        'availability; "No" indicates the feature is absent or in limited beta only.'
    )
    run = notes.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
