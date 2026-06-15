"""
Initial Setup: Writer document with empty footer
Task ID: writer_fs_079
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_079'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -- Title --
    title = doc.add_heading("Q1 2025 Project Status Report", level=0)

    # -- Subtitle --
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Prepared by the Project Management Office")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # -- Section 1: Executive Summary --
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report provides a comprehensive overview of project activities "
        "during Q1 2025. The quarter saw significant progress across all three "
        "major workstreams, with the infrastructure modernization initiative "
        "reaching its second milestone ahead of schedule. Client satisfaction "
        "scores improved by 12% compared to Q4 2024."
    )
    doc.add_paragraph(
        "Key highlights include the successful deployment of the automated "
        "testing framework, the onboarding of two new enterprise clients, and "
        "the completion of the security audit with zero critical findings."
    )

    # -- Section 2: Project Milestones --
    doc.add_heading("2. Project Milestones", level=1)
    doc.add_paragraph(
        "The following milestones were achieved during the reporting period:"
    )
    milestones = [
        "Infrastructure Modernization Phase 2 - Completed January 28, 2025",
        "Automated Testing Framework v2.1 - Deployed February 14, 2025",
        "Security Compliance Audit - Passed March 5, 2025",
        "Client Portal Redesign - Beta launch March 20, 2025",
    ]
    for m in milestones:
        doc.add_paragraph(m, style="List Bullet")

    # -- Section 3: Budget Overview --
    doc.add_heading("3. Budget Overview", level=1)
    doc.add_paragraph(
        "Total expenditure for Q1 2025 was $1,245,800 against a budgeted "
        "amount of $1,300,000, resulting in a favorable variance of $54,200. "
        "The engineering team accounted for 62% of total spend, followed by "
        "operations at 23% and marketing at 15%."
    )

    # Budget table
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Department", "Budget", "Actual", "Variance"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    budget_data = [
        ["Engineering", "$806,000", "$772,400", "+$33,600"],
        ["Operations", "$299,000", "$286,530", "+$12,470"],
        ["Marketing", "$195,000", "$186,870", "+$8,130"],
        ["Total", "$1,300,000", "$1,245,800", "+$54,200"],
    ]
    for r, row_data in enumerate(budget_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # -- Section 4: Risk Assessment --
    doc.add_heading("4. Risk Assessment", level=1)
    doc.add_paragraph(
        "Three medium-priority risks were identified during the quarter. "
        "The vendor dependency risk for the cloud migration has been mitigated "
        "through the establishment of a secondary provider agreement. Staff "
        "turnover in the DevOps team remains a concern, with two open positions "
        "currently in the final interview stage."
    )

    # -- Section 5: Next Steps --
    doc.add_heading("5. Next Steps for Q2 2025", level=1)
    next_steps = [
        "Complete Infrastructure Modernization Phase 3 by April 30",
        "Launch client portal to all enterprise accounts by May 15",
        "Begin mobile application development sprint cycle",
        "Conduct mid-year performance reviews for all project staff",
        "Submit annual compliance certification documentation",
    ]
    for i, step in enumerate(next_steps, 1):
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph(
        "For questions regarding this report, please contact the PMO at "
        "pmo@techsolutions.com or extension 4520."
    )

    # -- Footer: enabled but EMPTY --
    # Enable the footer by accessing it, but leave content empty
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear any default content
    for para in footer.paragraphs:
        para.text = ""

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
