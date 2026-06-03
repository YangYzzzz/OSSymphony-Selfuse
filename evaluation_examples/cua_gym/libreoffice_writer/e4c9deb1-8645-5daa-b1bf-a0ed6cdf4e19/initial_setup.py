"""
Initial Setup: bilingual_report.docx with French paragraph (no language tag)
Task ID: writer_edit_063
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_edit_063'
OUTPUT = f'{WORKDIR}/bilingual_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Paragraph 1: Introduction (English)
    p1 = doc.add_paragraph(
        "This annual report provides a comprehensive overview of our company's "
        "performance and strategic direction for the fiscal year 2024. We are pleased "
        "to present these findings to our shareholders and stakeholders."
    )

    # Paragraph 2: Key highlights (English)
    p2 = doc.add_paragraph(
        "Key highlights from this year include a 12% increase in overall revenue, "
        "successful expansion into three new markets, and the launch of our flagship "
        "product line. Our operational efficiency improved significantly across all "
        "departments, driven by streamlined processes and digital transformation initiatives."
    )

    # Paragraph 3: French translation (third paragraph — language NOT set to French)
    # This is the paragraph the agent must change to French language
    p3 = doc.add_paragraph(
        "Les résultats de cette année montrent une amélioration significative de la "
        "performance globale de l'entreprise."
    )
    # NOTE: Language is intentionally left as English (default) so the French words
    # are flagged as spelling errors. The agent must set this to French (France).

    # Paragraph 4: Outlook (English)
    p4 = doc.add_paragraph(
        "Looking ahead to the coming year, management anticipates continued growth "
        "supported by new product development, strategic partnerships, and an expanded "
        "customer base in the Asia-Pacific region. We remain committed to sustainable "
        "practices and delivering long-term value."
    )

    # Paragraph 5: Conclusion (English)
    p5 = doc.add_paragraph(
        "We would like to express our gratitude to all employees, partners, and "
        "customers for their contributions to another successful year. The dedication "
        "of our team has been instrumental in achieving these results."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
