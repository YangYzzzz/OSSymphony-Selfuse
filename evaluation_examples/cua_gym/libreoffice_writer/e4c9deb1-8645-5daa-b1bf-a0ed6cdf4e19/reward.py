"""
Reward Script: Change spelling language for French paragraph to French (France)
Task ID: writer_edit_063
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Third paragraph runs have language explicitly set to fr-FR
  Component 2 (0.3): Third paragraph paragraph-level rPr also carries fr-FR language marker
  Component 3 (0.2): English paragraphs (0,1,3,4) remain unchanged (no French language override)
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_edit_063'
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# The French paragraph text (paragraph index 2) to locate the target paragraph
FRENCH_TEXT_START = 'Les r'  # Unique prefix to identify the French paragraph


def get_run_lang(run):
    """Return the w:val of the first w:lang element found in the run's rPr, or None."""
    lang_elem = run._element.find(f'.//{{{WNS}}}lang')
    if lang_elem is not None:
        return lang_elem.get(f'{{{WNS}}}val')
    return None


def get_para_level_lang(para):
    """Return paragraph-level rPr language (pPr > rPr > lang w:val), or None."""
    pPr = para._element.find(f'{{{WNS}}}pPr')
    if pPr is None:
        return None
    pPr_rPr = pPr.find(f'{{{WNS}}}rPr')
    if pPr_rPr is None:
        return None
    lang_elem = pPr_rPr.find(f'{{{WNS}}}lang')
    if lang_elem is None:
        return None
    return lang_elem.get(f'{{{WNS}}}val')


def verify_task(file_path):
    """
    Verify that the French paragraph (paragraph 2) has its language changed to fr-FR,
    while the English paragraphs remain unchanged.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: verify the document has at least 5 paragraphs and para 2 is the French one
    if len(doc.paragraphs) < 5:
        print(f"CRITICAL: Expected at least 5 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    french_para = doc.paragraphs[2]
    if not french_para.text.startswith(FRENCH_TEXT_START):
        print(f"CRITICAL: Para 2 does not start with expected French text. Found: {repr(french_para.text[:50])}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: French paragraph runs have language set to fr-FR (0.5 points)
    # This FAILS on initial (no lang element) and PASSES on golden (fr-FR set)
    try:
        french_runs = french_para.runs
        if not french_runs:
            print("FAIL: Component 1 — French paragraph has no runs")
        else:
            run_langs = [get_run_lang(run) for run in french_runs]
            # Check that all non-None languages are fr-FR and at least one is set
            has_fr_fr = any(lang == 'fr-FR' for lang in run_langs)
            has_wrong_lang = any(lang is not None and lang != 'fr-FR' for lang in run_langs)

            if has_fr_fr and not has_wrong_lang:
                print(f"PASS: Component 1 — French paragraph runs have lang=fr-FR (run_langs={run_langs}) (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 1 — Expected fr-FR on French paragraph runs, found run_langs={run_langs}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraph-level rPr also carries fr-FR language marker (0.3 points)
    # In LibreOffice Writer, Format > Character > Language sets both pPr/rPr/lang AND run/rPr/lang
    # This FAILS on initial (no pPr lang) and PASSES on golden (fr-FR in pPr/rPr/lang)
    try:
        para_lang = get_para_level_lang(french_para)
        if para_lang == 'fr-FR':
            print(f"PASS: Component 2 — French paragraph has para-level lang=fr-FR (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected para-level lang=fr-FR, found para_lang={para_lang}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: English paragraphs (0, 1, 3, 4) do NOT have a French language override (0.2 points)
    # This ensures the task only changed the French paragraph, not the whole document
    # This PASSES on initial (no lang anywhere) AND on golden (only para 2 changed)
    # BUT we check it as part of the golden state validation: any new fr-FR on English paras is wrong
    # NOTE: For initial env, since no runs have lang set anywhere, this will also pass on initial.
    # We integrate it as a sub-condition to Component 1 to avoid awarding points in initial.
    # Instead, make it conditional on Component 1 having passed (i.e., only matters if fr-FR is present on para 2)
    try:
        english_para_indices = [0, 1, 3, 4]
        english_contaminated = []
        for idx in english_para_indices:
            para = doc.paragraphs[idx]
            # Check run-level lang
            for run in para.runs:
                lang = get_run_lang(run)
                if lang is not None and 'fr' in lang.lower():
                    english_contaminated.append((idx, 'run', lang))
            # Check para-level lang
            p_lang = get_para_level_lang(para)
            if p_lang is not None and 'fr' in p_lang.lower():
                english_contaminated.append((idx, 'para', p_lang))

        if not english_contaminated:
            # Only award these points if Component 1 also passed (otherwise initial would score 0.2)
            # Check that para 2 actually has fr-FR (meaning the task was done)
            french_run_langs = [get_run_lang(run) for run in french_para.runs]
            if any(lang == 'fr-FR' for lang in french_run_langs):
                print(f"PASS: Component 3 — English paragraphs remain unaffected (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Skipped (task not completed, no fr-FR on French paragraph)")
        else:
            print(f"FAIL: Component 3 — English paragraphs have unexpected French language: {english_contaminated}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/bilingual_report.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
