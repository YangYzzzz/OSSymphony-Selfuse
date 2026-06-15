"""
Reward Script: SOP-042 Equipment Calibration Procedure in LibreOffice Writer
Task ID: writer_wf_055
Domain: libreoffice_writer
Scoring:
  Component 1: Title with correct text and Title style (0.10)
  Component 2: Document info fields (0.10)
  Component 3: Seven Heading 2 sections (0.20)
  Component 4: Definitions table with 4 terms (0.15)
  Component 5: Five bulleted equipment items (0.10)
  Component 6: Eight numbered procedure steps (0.15)
  Component 7: Three quality check verification points (0.10)
  Component 8: Revision history table at end (0.10)
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_055'


def persist_app_state(domain: str):
    """Try to save any unsaved LibreOffice edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify SOP document creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    full_text = '\n'.join(p.text for p in paragraphs)

    # Quick gate: if document is essentially empty, return 0
    if len(paragraphs) < 5:
        print(f"FAIL: Document has only {len(paragraphs)} paragraphs - appears empty/minimal")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title with correct text and Title style (0.10 points)
    try:
        title_found = False
        for p in paragraphs:
            if 'SOP-042' in p.text and 'Equipment Calibration' in p.text:
                style_name = p.style.name if p.style else ''
                if 'Title' in style_name or 'Heading 1' in style_name or style_name == 'Title':
                    print(f"PASS: Component 1 — Title found with style '{style_name}' (0.10 pts)")
                    total_score += 0.10
                    title_found = True
                else:
                    # Partial: text is right but style is wrong
                    print(f"PARTIAL: Component 1 — Title text found but style is '{style_name}', expected Title/Heading 1 (0.05 pts)")
                    total_score += 0.05
                    title_found = True
                break
        if not title_found:
            print(f"FAIL: Component 1 — Title 'SOP-042: Equipment Calibration Procedure' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Document info fields (0.10 points)
    # Must have: Effective Date, Department, Author, Approved By
    try:
        info_fields = ['Effective Date', 'Department', 'Author', 'Approved By']
        found_count = 0
        for field in info_fields:
            for p in paragraphs:
                if field.lower() in p.text.lower():
                    found_count += 1
                    break
        if found_count == 4:
            print(f"PASS: Component 2 — All 4 document info fields found (0.10 pts)")
            total_score += 0.10
        elif found_count >= 2:
            partial = round(0.10 * found_count / 4, 2)
            print(f"PARTIAL: Component 2 — {found_count}/4 info fields found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {found_count}/4 document info fields found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Seven Heading 2 sections (0.20 points)
    # Expected sections: Purpose, Scope, Definitions, Required Equipment,
    #   Procedure Steps, Quality Checks, Records and Documentation
    try:
        expected_sections = [
            'Purpose', 'Scope', 'Definitions', 'Required Equipment',
            'Procedure', 'Quality', 'Records'
        ]
        heading2_paras = [p for p in paragraphs if p.style and 'Heading 2' in p.style.name]
        heading2_texts = [p.text.lower() for p in heading2_paras]

        matched_sections = 0
        for section in expected_sections:
            for h_text in heading2_texts:
                if section.lower() in h_text:
                    matched_sections += 1
                    break

        if matched_sections >= 7:
            print(f"PASS: Component 3 — All 7 Heading 2 sections found ({len(heading2_paras)} total H2 paras) (0.20 pts)")
            total_score += 0.20
        elif matched_sections >= 4:
            partial = round(0.20 * matched_sections / 7, 2)
            print(f"PARTIAL: Component 3 — {matched_sections}/7 expected sections found as Heading 2 ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {matched_sections}/7 expected Heading 2 sections found. H2 headings: {heading2_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Definitions table with 4 terms (0.15 points)
    # Table should have header row (Term, Definition) + at least 4 data rows
    try:
        def_table_found = False
        for table in doc.tables:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'term' in headers and 'definition' in headers:
                data_rows = len(table.rows) - 1  # exclude header
                if data_rows >= 4:
                    print(f"PASS: Component 4 — Definitions table found with {data_rows} term rows (0.15 pts)")
                    total_score += 0.15
                    def_table_found = True
                elif data_rows >= 2:
                    partial = round(0.15 * data_rows / 4, 2)
                    print(f"PARTIAL: Component 4 — Definitions table found but only {data_rows}/4 terms ({partial} pts)")
                    total_score += partial
                    def_table_found = True
                else:
                    print(f"FAIL: Component 4 — Definitions table found but only {data_rows} term rows")
                    def_table_found = True
                break
        if not def_table_found:
            print(f"FAIL: Component 4 — No definitions table with 'Term'/'Definition' headers found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Five bulleted equipment items (0.10 points)
    try:
        bullet_paras = [p for p in paragraphs if p.style and 'List Bullet' in p.style.name and p.text.strip()]
        if len(bullet_paras) >= 5:
            print(f"PASS: Component 5 — {len(bullet_paras)} bulleted items found (0.10 pts)")
            total_score += 0.10
        elif len(bullet_paras) >= 3:
            partial = round(0.10 * len(bullet_paras) / 5, 2)
            print(f"PARTIAL: Component 5 — {len(bullet_paras)}/5 bulleted items found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {len(bullet_paras)} List Bullet paragraphs found, need 5")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Eight numbered procedure steps (0.15 points)
    try:
        numbered_paras = [p for p in paragraphs if p.style and 'List Number' in p.style.name and p.text.strip()]
        if len(numbered_paras) >= 8:
            print(f"PASS: Component 6 — {len(numbered_paras)} numbered steps found (0.15 pts)")
            total_score += 0.15
        elif len(numbered_paras) >= 4:
            partial = round(0.15 * len(numbered_paras) / 8, 2)
            print(f"PARTIAL: Component 6 — {len(numbered_paras)}/8 numbered steps found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — Only {len(numbered_paras)} List Number paragraphs found, need 8")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Three quality check verification points (0.10 points)
    # Look for paragraphs mentioning "verification point" after the Quality Checks heading
    try:
        qc_count = 0
        in_qc_section = False
        for p in paragraphs:
            if p.style and 'Heading 2' in p.style.name and 'quality' in p.text.lower():
                in_qc_section = True
                continue
            if in_qc_section and p.style and 'Heading 2' in p.style.name:
                # Moved to next section
                break
            if in_qc_section and p.text.strip():
                # Count non-empty paragraphs in Quality Checks section as verification points
                qc_count += 1

        if qc_count >= 3:
            print(f"PASS: Component 7 — {qc_count} quality check verification points found (0.10 pts)")
            total_score += 0.10
        elif qc_count >= 1:
            partial = round(0.10 * qc_count / 3, 2)
            print(f"PARTIAL: Component 7 — {qc_count}/3 quality check points found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 7 — No quality check verification points found in Quality Checks section")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: Revision history table at end (0.10 points)
    # Should be a table with columns like Revision, Date, Author, Description
    try:
        rev_table_found = False
        # Check last table or any table with revision-like headers
        for table in doc.tables:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if any('revision' in h or 'rev' in h for h in headers) and \
               any('date' in h for h in headers):
                data_rows = len(table.rows) - 1
                if data_rows >= 1:
                    print(f"PASS: Component 8 — Revision history table found with {data_rows} entries (0.10 pts)")
                    total_score += 0.10
                    rev_table_found = True
                else:
                    print(f"FAIL: Component 8 — Revision table found but no data rows")
                    rev_table_found = True
                break
        if not rev_table_found:
            # Also check if "revision history" appears in the text and a table follows
            if 'revision history' in full_text.lower() and len(doc.tables) >= 2:
                last_table = doc.tables[-1]
                if len(last_table.rows) >= 2:
                    print(f"PASS: Component 8 — Revision history table found (last table, {len(last_table.rows)} rows) (0.10 pts)")
                    total_score += 0.10
                    rev_table_found = True
        if not rev_table_found:
            print(f"FAIL: Component 8 — No revision history table found")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
