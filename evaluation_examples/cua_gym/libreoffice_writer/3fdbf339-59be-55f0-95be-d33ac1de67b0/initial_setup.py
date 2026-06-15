"""
Initial Setup: Insert a fixed date field in a Writer document
Task ID: writer_tm_080
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_080'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Service Agreement', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Parties Section ---
    doc.add_heading('Parties', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('This Service Agreement ("Agreement") is entered into by and between ')
    run_company = p1.add_run('Meridian Solutions Inc.')
    run_company.bold = True
    p1.add_run(' ("Provider"), located at 450 Technology Drive, Suite 300, San Jose, CA 95110, ')
    p1.add_run('and ')
    run_client = p1.add_run('Northwind Enterprises LLC')
    run_client.bold = True
    p1.add_run(' ("Client"), located at 1200 Commerce Boulevard, Austin, TX 78701.')

    # --- Scope of Services ---
    doc.add_heading('Scope of Services', level=1)
    p2 = doc.add_paragraph(
        'The Provider agrees to deliver the following managed IT services to the Client:'
    )
    services = [
        'Network infrastructure monitoring and maintenance (24/7)',
        'Cloud hosting and management (AWS / Azure)',
        'Cybersecurity threat detection and incident response',
        'Quarterly technology audits and compliance reviews',
        'End-user technical support (Tier 1 through Tier 3)',
    ]
    for svc in services:
        doc.add_paragraph(svc, style='List Bullet')

    # --- Effective Date paragraph ---
    doc.add_heading('Term and Duration', level=1)

    # The key paragraph: "Effective Date: " with nothing after it
    p_date = doc.add_paragraph()
    run_label = p_date.add_run('Effective Date: ')
    run_label.bold = True

    p_term = doc.add_paragraph(
        'This Agreement shall remain in effect for a period of twenty-four (24) months '
        'from the Effective Date, unless terminated earlier in accordance with the '
        'provisions set forth herein.'
    )

    # --- Compensation ---
    doc.add_heading('Compensation', level=1)
    doc.add_paragraph(
        'The Client agrees to pay the Provider a monthly service fee of $12,500.00 USD, '
        'due on the first business day of each calendar month. Late payments shall incur '
        'a penalty of 1.5% per month on the outstanding balance.'
    )

    # --- Confidentiality ---
    doc.add_heading('Confidentiality', level=1)
    doc.add_paragraph(
        'Both parties agree to maintain the confidentiality of all proprietary information, '
        'trade secrets, and business data exchanged during the term of this Agreement. '
        'This obligation shall survive termination of the Agreement for a period of '
        'three (3) years.'
    )

    # --- Termination ---
    doc.add_heading('Termination', level=1)
    doc.add_paragraph(
        'Either party may terminate this Agreement with sixty (60) days written notice. '
        'In the event of a material breach, the non-breaching party may terminate '
        'immediately upon written notice, provided that the breaching party has been '
        'given fifteen (15) business days to cure the breach.'
    )

    # --- Signatures ---
    doc.add_heading('Signatures', level=1)

    # Create signature table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    table.cell(0, 0).text = 'Provider: Meridian Solutions Inc.'
    table.cell(0, 1).text = 'Client: Northwind Enterprises LLC'

    table.cell(1, 0).text = 'Name: Elena Rodriguez, CEO'
    table.cell(1, 1).text = 'Name: David Park, VP Operations'

    table.cell(2, 0).text = 'Signature: _______________________'
    table.cell(2, 1).text = 'Signature: _______________________'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
