"""
FINAL REWARD SCRIPT - SUCCESS
Task: In my document, paragraphs 2, 3, and 4 have ragged right edges and words spill over awkwardly. How do I make just those three paragraphs fully justified and turn hyphenation on for them so the lines even out?
Generated: 2025-09-10 19:58:51
Status: success
Model: azure-o3
Total Steps: 4
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

"""
Reward script for verifying the task:
"In my document, paragraphs 2, 3, and 4 have ragged right edges and words spill over awkwardly. 
How do I make just those three paragraphs fully justified and turn hyphenation on for them so the lines even out?"

This script checks the provided DOCX file to ensure that:
1. Paragraphs 2, 3, and 4 (0-based indices 1, 2, 3) are fully-justified.
2. Automatic hyphenation is enabled (not suppressed) for those same paragraphs.

Scoring (progressive):
 • Up to 0.5 points for correct justification (0.1667 each paragraph).
 • Up to 0.5 points for correct hyphenation settings (0.1667 each paragraph).
The script prints detailed diagnostics and finishes with "REWARD: X.X" where X.X ∈ [0.0, 1.0].
"""

FILE_PATH = "/home/user/in_my_document_paragraphs_2_3_and_4_have_ragged_right_edges_and_words_spill_over_awkwardly_how_do_i_.docx"

def is_hyphenation_on(paragraph):
    """Return True if automatic hyphenation is enabled for the paragraph."""
    pPr = paragraph._p.pPr
    # If no paragraph properties, hyphenation is ON by default (not suppressed)
    if pPr is None:
        return True

    # Word suppresses hyphenation via <w:suppressAutoHyphens/> element.
    suppress_elems = pPr.findall(qn("w:suppressAutoHyphens"))
    if not suppress_elems:
        return True  # not suppressed → hyphenation allowed

    # Element present: check its value attribute
    val = suppress_elems[0].get(qn("w:val"))
    # If val missing or truthy ("1", "true", "on") then suppression ON → hyphenation OFF
    if val is None or val in {"1", "true", "on"}:
        return False
    # val "0"/"false" disables suppression → hyphenation ON
    return True

def is_fully_justified(paragraph):
    """Return True if paragraph alignment is any justification variant."""
    align = paragraph.paragraph_format.alignment
    # Accept common justification enums
    justify_values = {
        WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE,
        getattr(WD_PARAGRAPH_ALIGNMENT, "JUSTIFY_HIGH", WD_PARAGRAPH_ALIGNMENT.JUSTIFY),
        getattr(WD_PARAGRAPH_ALIGNMENT, "JUSTIFY_MED", WD_PARAGRAPH_ALIGNMENT.JUSTIFY),
        getattr(WD_PARAGRAPH_ALIGNMENT, "JUSTIFY_LOW", WD_PARAGRAPH_ALIGNMENT.JUSTIFY),
    }
    return align in justify_values

def verify_task(file_path):
    print(f"Verifying task for document: {file_path}")

    if not os.path.exists(file_path):
        print("✗ File not found.")
        return 0.0

    try:
        doc = Document(file_path)
        print(f"✓ Document loaded. Total paragraphs: {len(doc.paragraphs)}")
    except Exception as e:
        print(f"✗ Error loading document: {e}")
        return 0.0

    if len(doc.paragraphs) < 4:
        print("✗ Document contains fewer than 4 paragraphs – cannot verify.")
        return 0.0

    target_indices = [1, 2, 3]  # paragraphs 2-4 (0-based)
    justified_ok = 0
    hyphen_ok = 0

    for idx in target_indices:
        para = doc.paragraphs[idx]
        text_preview = para.text.strip()[:60]
        justified = is_fully_justified(para)
        hyphen = is_hyphenation_on(para)

        print(f"Paragraph {idx+1}: '{text_preview}...' if longer")
        print(f"  - Fully justified: {justified}")
        print(f"  - Hyphenation ON : {hyphen}")

        if justified:
            justified_ok += 1
        if hyphen:
            hyphen_ok += 1

    # Progressive scoring
    alignment_score = (justified_ok / 3) * 0.5  # up to 0.5
    hyphen_score = (hyphen_ok / 3) * 0.5        # up to 0.5
    total_score = round(alignment_score + hyphen_score, 4)

    print(f"Alignment score : {alignment_score:.2f}/0.50")
    print(f"Hyphenation score: {hyphen_score:.2f}/0.50")
    print(f"REWARD: {total_score}")

    return total_score

# Execute verification when script is run
if __name__ == "__main__":
    verify_task(FILE_PATH)
