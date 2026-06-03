"""
Reward Script: Search for all italic text and replace formatting with bold italic
Task ID: writer_frd_006
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): No italic-only runs remain (all converted)
  Component 2 (0.4): All 15 expected text segments are bold-italic
  Component 3 (0.2): Text content is preserved (unchanged from initial)
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_006'

# The 15 expected italic text segments from the task context
EXPECTED_ITALIC_TEXTS = [
    "Cognitive Load During Problem Solving",
    "Multimedia Learning",
    "The Magical Number Seven, Plus or Minus Two",
    "sine qua non",
    "a priori",
    "Working Memory",
    "The Architecture of Cognition",
    "in vivo",
    "Learning from Worked-Out Examples",
    "per se",
    "et al.",
    "Interactive Multimodal Learning Environments",
    "The Journal of Educational Psychology",
    "de facto",
    "Why Minimal Guidance During Instruction Does Not Work",
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all runs with their formatting
    italic_only_runs = []
    bold_italic_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            if run.font.italic and not run.font.bold:
                italic_only_runs.append(text)
            elif run.font.italic and run.font.bold:
                bold_italic_runs.append(text)

    # Component 1: No italic-only runs remain (0.4 points)
    # In initial_env there are 15 italic-only runs; in golden there should be 0
    try:
        if len(italic_only_runs) == 0:
            print(f"PASS: Component 1 — No italic-only runs remain (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Found {len(italic_only_runs)} italic-only runs still remaining")
            for t in italic_only_runs[:5]:
                print(f"       - '{t}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 15 expected segments are now bold-italic (0.4 points)
    # Award proportional credit based on how many of the 15 are found as bold-italic
    try:
        found_count = 0
        for expected in EXPECTED_ITALIC_TEXTS:
            # Check if any bold-italic run contains this text
            if any(expected in bi_text or bi_text in expected for bi_text in bold_italic_runs):
                found_count += 1
            else:
                print(f"  MISS: Bold-italic not found for: '{expected}'")

        if found_count == len(EXPECTED_ITALIC_TEXTS):
            print(f"PASS: Component 2 — All {found_count}/15 segments are bold-italic (0.4 pts)")
            total_score += 0.4
        elif found_count > 0:
            partial = 0.4 * (found_count / len(EXPECTED_ITALIC_TEXTS))
            print(f"PARTIAL: Component 2 — {found_count}/15 segments are bold-italic ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No expected segments found as bold-italic")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Bold-italic runs preserve original text AND no italic-only remain (0.2 points)
    # This is a compound check: verifies that the formatting change was done correctly
    # (bold added to italic runs) without altering text content. It only passes when
    # the task change has been applied (bold-italic runs exist) AND text is intact.
    try:
        full_text = "\n".join(para.text for para in doc.paragraphs)
        all_text_present = all(expected in full_text for expected in EXPECTED_ITALIC_TEXTS)

        # This component requires BOTH: task was done (bold-italic runs exist)
        # AND text was preserved. On initial_env, bold_italic_runs is empty so this fails.
        if len(bold_italic_runs) >= 15 and all_text_present and len(italic_only_runs) == 0:
            print(f"PASS: Component 3 — Formatting converted AND text preserved (0.2 pts)")
            total_score += 0.2
        elif len(bold_italic_runs) > 0 and all_text_present:
            print(f"PARTIAL: Component 3 — Some bold-italic present and text preserved (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 3 — bold_italic_runs={len(bold_italic_runs)}, text_present={all_text_present}, italic_only={len(italic_only_runs)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
