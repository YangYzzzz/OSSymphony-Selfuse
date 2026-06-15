"""
Initial Setup: Research paper with italic-formatted text segments
Task ID: writer_frd_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_normal_run(para, text, font_name="Times New Roman", font_size=12):
    """Add a normal (non-italic) run."""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    return run


def add_italic_run(para, text, font_name="Times New Roman", font_size=12):
    """Add an italic run."""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.italic = True
    return run


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ---- Title ----
    title = doc.add_heading("A Comprehensive Literature Review on Cognitive Load Theory\nin Educational Technology", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = author.add_run("Dr. Elena Vasquez, Department of Educational Psychology\nUniversity of Western Ontario\nRevised March 2025")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # ---- Abstract ----
    doc.add_heading("Abstract", level=1)
    abs_para = doc.add_paragraph()
    add_normal_run(abs_para, "This literature review examines the evolution of Cognitive Load Theory (CLT) and its applications in modern educational technology. Drawing from seminal works including Sweller's ")
    add_italic_run(abs_para, "Cognitive Load During Problem Solving")  # ITALIC #1
    add_normal_run(abs_para, " (1988) and Mayer's ")
    add_italic_run(abs_para, "Multimedia Learning")  # ITALIC #2
    add_normal_run(abs_para, " (2009), this review synthesizes over four decades of research. We analyze the theoretical underpinnings, empirical evidence, and practical implications of CLT across diverse educational contexts. Our findings suggest that the integration of CLT principles into instructional design significantly enhances learning outcomes, particularly when cognitive demands are carefully calibrated to learner expertise levels.")

    # ---- 1. Introduction ----
    doc.add_heading("1. Introduction", level=1)
    p1 = doc.add_paragraph()
    add_normal_run(p1, "The study of human cognition in educational settings has a long and rich history. Since the publication of Miller's landmark paper ")
    add_italic_run(p1, "The Magical Number Seven, Plus or Minus Two")  # ITALIC #3
    add_normal_run(p1, " in 1956, researchers have sought to understand the limitations and capacities of working memory. This foundational work established that human information processing capacity is inherently constrained, a principle that has profound implications for how we design learning environments and instructional materials.")

    p2 = doc.add_paragraph()
    add_normal_run(p2, "Cognitive Load Theory, first proposed by John Sweller in the late 1980s, provides a comprehensive framework for understanding how instructional design interacts with human cognitive architecture. The theory distinguishes between three types of cognitive load: intrinsic, extraneous, and germane. As Paas and van Merri\u00ebnboer noted in their influential review, the management of these load types is ")
    add_italic_run(p2, "sine qua non")  # ITALIC #4
    add_normal_run(p2, " for effective instruction.")

    p3 = doc.add_paragraph()
    add_normal_run(p3, "The rapid advancement of digital technologies has created both opportunities and challenges for educators. Tablet-based learning, virtual reality simulations, and adaptive tutoring systems all impose varying cognitive demands on learners. Understanding these demands ")
    add_italic_run(p3, "a priori")  # ITALIC #5
    add_normal_run(p3, " allows instructional designers to create more effective learning experiences that respect the boundaries of human working memory.")

    # ---- 2. Theoretical Foundations ----
    doc.add_heading("2. Theoretical Foundations", level=1)
    p4 = doc.add_paragraph()
    add_normal_run(p4, "The theoretical foundations of CLT rest on several well-established principles from cognitive psychology. Baddeley and Hitch's model of working memory, presented in ")
    add_italic_run(p4, "Working Memory")  # ITALIC #6
    add_normal_run(p4, " (1974), provides the architectural basis for understanding capacity limitations. Their multi-component model, comprising the central executive, phonological loop, visuospatial sketchpad, and episodic buffer, remains the most widely accepted framework for describing short-term information processing.")

    p5 = doc.add_paragraph()
    add_normal_run(p5, "Schema theory further enriches our understanding. As articulated by Anderson in ")
    add_italic_run(p5, "The Architecture of Cognition")  # ITALIC #7
    add_normal_run(p5, " (1983), knowledge is organized into interconnected structures called schemas. These schemas allow experts to process complex information efficiently by chunking related elements into single units. The transition from novice to expert can be characterized as the progressive development of increasingly sophisticated schemas that reduce the effective cognitive load of domain-specific tasks.")

    p6 = doc.add_paragraph()
    add_normal_run(p6, "It is worth noting that the relationship between expertise and cognitive load operates ")
    add_italic_run(p6, "in vivo")  # ITALIC #8
    add_normal_run(p6, " in complex and often unpredictable ways. Laboratory studies, while valuable, do not always capture the full range of factors that influence cognitive processing in authentic learning environments. Ecological validity remains a persistent concern in CLT research.")

    # ---- 3. Empirical Evidence ----
    doc.add_heading("3. Empirical Evidence", level=1)

    doc.add_heading("3.1 The Worked Example Effect", level=2)
    p7 = doc.add_paragraph()
    add_normal_run(p7, "Among the most robust findings in CLT research is the worked example effect. Sweller and Cooper (1985) demonstrated that students who studied worked examples outperformed those who solved equivalent problems independently. This effect has been replicated across numerous domains, from mathematics to computer programming. The theoretical explanation, as detailed in Renkl's ")
    add_italic_run(p7, "Learning from Worked-Out Examples")  # ITALIC #9
    add_normal_run(p7, " (2014), centers on the reduction of extraneous cognitive load: worked examples eliminate the need for means-ends analysis, freeing working memory resources for schema construction.")

    doc.add_heading("3.2 The Split-Attention Effect", level=2)
    p8 = doc.add_paragraph()
    add_normal_run(p8, "When learners must mentally integrate multiple sources of information that are physically separated, the resulting split-attention effect imposes unnecessary extraneous load. Chandler and Sweller (1992) provided compelling evidence that integrated formats, where text and diagrams are physically combined, significantly improve learning outcomes compared to split-source formats. This finding has been extensively validated in domains ranging from electrical engineering to medical education, and its implications ")
    add_italic_run(p8, "per se")  # ITALIC #10
    add_normal_run(p8, " extend far beyond traditional classroom settings into the design of software interfaces and online learning platforms.")

    doc.add_heading("3.3 The Redundancy Effect", level=2)
    p9 = doc.add_paragraph()
    add_normal_run(p9, "The redundancy effect occurs when multiple sources of information that are self-contained are presented simultaneously, forcing learners to process redundant material. Kalyuga ")
    add_italic_run(p9, "et al.")  # ITALIC #11
    add_normal_run(p9, " (2003) demonstrated that removing redundant information can significantly improve learning outcomes, particularly for more experienced learners. This finding challenges the common instructional practice of providing multiple representations of the same content under the assumption that more information is always beneficial.")

    # ---- 4. Applications in Educational Technology ----
    doc.add_heading("4. Applications in Educational Technology", level=1)
    p10 = doc.add_paragraph()
    add_normal_run(p10, "The translation of CLT principles into practical educational technology has been a major focus of research over the past two decades. Moreno and Mayer's ")
    add_italic_run(p10, "Interactive Multimodal Learning Environments")  # ITALIC #12
    add_normal_run(p10, " (2007) established a framework for designing multimedia learning systems that manage cognitive load effectively. Their research demonstrated that the modality principle, which suggests presenting information across visual and auditory channels, reduces cognitive load by distributing processing demands across multiple working memory subsystems.")

    p11 = doc.add_paragraph()
    add_normal_run(p11, "Adaptive learning technologies represent a particularly promising application of CLT. These systems dynamically adjust content difficulty based on learner performance, effectively managing intrinsic cognitive load in real time. Research by Kalyuga and Sweller, published in ")
    add_italic_run(p11, "The Journal of Educational Psychology")  # ITALIC #13
    add_normal_run(p11, " (2005), showed that the expertise reversal effect, where techniques beneficial for novices become detrimental for experts, can be mitigated through adaptive sequencing algorithms that continuously assess learner expertise and modify instructional support accordingly.")

    # ---- 5. Discussion ----
    doc.add_heading("5. Discussion and Future Directions", level=1)
    p12 = doc.add_paragraph()
    add_normal_run(p12, "Despite the substantial progress outlined in this review, several challenges remain. First, the measurement of cognitive load continues to be a ")
    add_italic_run(p12, "de facto")  # ITALIC #14
    add_normal_run(p12, " methodological obstacle. Subjective rating scales, physiological measures such as pupil dilation and heart rate variability, and dual-task paradigms each capture different aspects of cognitive load, yet no single measure provides a comprehensive assessment. Developing more reliable and valid measurement instruments remains a critical priority for the field.")

    p13 = doc.add_paragraph()
    add_normal_run(p13, "Second, the increasing diversity of learning technologies demands continued investigation into how CLT principles apply across different modalities and platforms. As Kirschner and colleagues argued in their provocative paper ")
    add_italic_run(p13, "Why Minimal Guidance During Instruction Does Not Work")  # ITALIC #15
    add_normal_run(p13, " (2006), the failure to account for cognitive load limitations has contributed to the disappointing outcomes of many discovery-based and constructivist approaches to instruction. Future research should focus on developing more nuanced models that account for individual differences in working memory capacity, prior knowledge, and motivation.")

    p14 = doc.add_paragraph()
    add_normal_run(p14, "In conclusion, Cognitive Load Theory has established itself as one of the most influential frameworks in educational psychology. Its principles provide concrete, actionable guidance for the design of instructional materials and learning technologies. As educational environments become increasingly complex and technology-mediated, the systematic application of CLT principles will be essential for ensuring that these environments support, rather than overwhelm, the cognitive capacities of learners.")

    # ---- References ----
    doc.add_heading("References", level=1)
    references = [
        "Anderson, J. R. (1983). The Architecture of Cognition. Harvard University Press.",
        "Baddeley, A. D., & Hitch, G. (1974). Working Memory. In G. H. Bower (Ed.), The Psychology of Learning and Motivation (Vol. 8, pp. 47\u201389). Academic Press.",
        "Chandler, P., & Sweller, J. (1992). The split-attention effect as a factor in the design of instruction. British Journal of Educational Psychology, 62(2), 233\u2013246.",
        "Kalyuga, S., Chandler, P., & Sweller, J. (2003). The expertise reversal effect. Educational Psychologist, 38(1), 23\u201331.",
        "Kalyuga, S., & Sweller, J. (2005). Rapid dynamic assessment of expertise. Journal of Educational Psychology, 97(1), 106\u2013114.",
        "Kirschner, P. A., Sweller, J., & Clark, R. E. (2006). Why minimal guidance during instruction does not work. Educational Psychologist, 41(2), 75\u201386.",
        "Mayer, R. E. (2009). Multimedia Learning (2nd ed.). Cambridge University Press.",
        "Miller, G. A. (1956). The magical number seven, plus or minus two. Psychological Review, 63(2), 81\u201397.",
        "Moreno, R., & Mayer, R. (2007). Interactive multimodal learning environments. Educational Psychology Review, 19(3), 309\u2013326.",
        "Renkl, A. (2014). Learning from worked-out examples. In R. K. Sawyer (Ed.), The Cambridge Handbook of the Learning Sciences (2nd ed., pp. 391\u2013412). Cambridge University Press.",
        "Sweller, J. (1988). Cognitive load during problem solving. Cognitive Science, 12(2), 257\u2013285.",
        "Sweller, J., & Cooper, G. A. (1985). The use of worked examples as a substitute for problem solving. Cognition and Instruction, 2(1), 59\u201389.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref, style="List Number")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
