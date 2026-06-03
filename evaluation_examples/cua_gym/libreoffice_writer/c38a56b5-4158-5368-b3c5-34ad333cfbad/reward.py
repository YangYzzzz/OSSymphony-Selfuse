"""
Reward Script: Mail merge letter with personalized case update fields
Task ID: writer_mt_021
Domain: libreoffice_writer
Scoring:
  Component 1: ClientName MERGEFIELD present in document (0.25 pts)
  Component 2: CaseNumber MERGEFIELD present in document (0.25 pts)
  Component 3: CaseStatus MERGEFIELD present in document (0.25 pts)
  Component 4: NextHearingDate MERGEFIELD present in document (0.25 pts)
"""

import os
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_021'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': W_NS}


def persist_app_state(domain: str):
    """Best-effort save in case LibreOffice has unsaved changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def extract_merge_fields(doc_path):
    """
    Extract all MERGEFIELD names from a .docx file by parsing the XML directly.
    Returns a set of field names found (e.g. {'ClientName', 'CaseNumber'}).
    """
    from docx import Document
    doc = Document(doc_path)
    merge_fields = set()

    # Iterate all paragraphs and look for instrText containing MERGEFIELD
    for para in doc.paragraphs:
        for instr in para._element.findall('.//w:instrText', NS):
            text = instr.text
            if text and 'MERGEFIELD' in text:
                # Extract field name: " MERGEFIELD ClientName " -> "ClientName"
                parts = text.strip().split()
                if len(parts) >= 2 and parts[0] == 'MERGEFIELD':
                    field_name = parts[1]
                    merge_fields.add(field_name)

    return merge_fields


def check_field_in_context(doc_path, field_name, context_text):
    """
    Check that a MERGEFIELD with the given field_name appears in a paragraph
    whose full text contains context_text.
    Returns True if found.
    """
    from docx import Document
    doc = Document(doc_path)

    for para in doc.paragraphs:
        if context_text not in para.text:
            continue
        # Found a paragraph with the expected context; check for the merge field
        for instr in para._element.findall('.//w:instrText', NS):
            text = instr.text
            if text and 'MERGEFIELD' in text:
                parts = text.strip().split()
                if len(parts) >= 2 and parts[1] == field_name:
                    return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires inserting 4 mail merge fields:
    ClientName, CaseNumber, CaseStatus, NextHearingDate
    into a law firm case update letter template.
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract all merge fields from the document
    try:
        all_fields = extract_merge_fields(file_path)
        print(f"INFO: Found merge fields: {all_fields}")
    except Exception as e:
        print(f"CRITICAL: Cannot parse merge fields: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: ClientName MERGEFIELD present in greeting (0.25 points)
    # Task requires "Dear <ClientName>," as the greeting
    try:
        if 'ClientName' in all_fields and check_field_in_context(file_path, 'ClientName', 'Dear'):
            print(f"PASS: Component 1 - ClientName MERGEFIELD found in greeting (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - ClientName MERGEFIELD not found in greeting context")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: CaseNumber MERGEFIELD present (0.25 points)
    # Task requires "Case #<CaseNumber>" in the body
    try:
        if 'CaseNumber' in all_fields and check_field_in_context(file_path, 'CaseNumber', 'Case #'):
            print(f"PASS: Component 2 - CaseNumber MERGEFIELD found in 'Case #' context (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 - CaseNumber MERGEFIELD not found in 'Case #' context")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: CaseStatus MERGEFIELD present (0.25 points)
    # Task requires "Your case status is currently: <CaseStatus>."
    try:
        if 'CaseStatus' in all_fields and check_field_in_context(file_path, 'CaseStatus', 'case status is currently'):
            print(f"PASS: Component 3 - CaseStatus MERGEFIELD found in status context (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 - CaseStatus MERGEFIELD not found in status context")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: NextHearingDate MERGEFIELD present (0.25 points)
    # Task requires "Your next hearing is scheduled for <NextHearingDate>."
    try:
        if 'NextHearingDate' in all_fields and check_field_in_context(file_path, 'NextHearingDate', 'next hearing is scheduled for'):
            print(f"PASS: Component 4 - NextHearingDate MERGEFIELD found in hearing date context (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 - NextHearingDate MERGEFIELD not found in hearing date context")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
