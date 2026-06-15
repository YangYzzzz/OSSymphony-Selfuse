"""
Initial Setup: Mail merge letter template for law firm case updates
Task ID: writer_mt_021
Domain: libreoffice_writer

Creates a professional legal case update letter template with plain-text
placeholders (NOT mail merge fields). The task requires the agent to
convert these into proper mail merge fields.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_021'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Law Firm Letterhead --
    heading = doc.add_heading('', level=0)
    run = heading.add_run('MORRISON, CHEN & ASSOCIATES')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = subtitle.add_run('Attorneys at Law')
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    address = doc.add_paragraph()
    address.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_addr = address.add_run(
        '1200 Pacific Avenue, Suite 400  |  San Francisco, CA 94115\n'
        'Phone: (415) 555-0182  |  Fax: (415) 555-0183  |  www.morrisonchen.com'
    )
    run_addr.font.size = Pt(9)
    run_addr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # -- Horizontal rule --
    doc.add_paragraph('_' * 72)

    # -- Date --
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(18)
    date_run = date_para.add_run('March 28, 2026')
    date_run.font.size = Pt(11)

    # -- Recipient placeholder --
    doc.add_paragraph('')  # blank line
    recipient = doc.add_paragraph()
    r1 = recipient.add_run('[Client Name]')
    r1.font.size = Pt(11)

    addr_line = doc.add_paragraph()
    a1 = addr_line.add_run('[Client Address]')
    a1.font.size = Pt(11)

    # -- Greeting with placeholder --
    doc.add_paragraph('')
    greeting = doc.add_paragraph()
    g_run = greeting.add_run('Dear [Client Name],')
    g_run.font.size = Pt(11)

    # -- Body paragraphs with placeholders --
    doc.add_paragraph('')
    body1 = doc.add_paragraph()
    b1_run = body1.add_run(
        'We are writing to provide you with an update regarding your legal matter. '
        'Our records indicate that your case, Case #[Case Number], has been actively '
        'reviewed by our legal team.'
    )
    b1_run.font.size = Pt(11)

    body2 = doc.add_paragraph()
    b2_run = body2.add_run(
        'Your case status is currently: [Case Status]. Our attorneys have been working '
        'diligently to ensure the best possible outcome for your matter. We want to keep '
        'you fully informed of all developments as they occur.'
    )
    b2_run.font.size = Pt(11)

    body3 = doc.add_paragraph()
    b3_run = body3.add_run(
        'Your next hearing is scheduled for [Next Hearing Date]. Please ensure you are '
        'available on this date and arrive at least 30 minutes prior to the scheduled time. '
        'If you have any conflicts with this date, please contact our office immediately so '
        'we can discuss options.'
    )
    b3_run.font.size = Pt(11)

    body4 = doc.add_paragraph()
    b4_run = body4.add_run(
        'In preparation for the hearing, we recommend gathering any relevant documentation '
        'and reviewing the case materials we previously provided. Our team will schedule a '
        'pre-hearing consultation with you in the coming days to review our strategy.'
    )
    b4_run.font.size = Pt(11)

    # -- Closing --
    doc.add_paragraph('')
    closing = doc.add_paragraph()
    c_run = closing.add_run('Should you have any questions or require further clarification, '
                            'please do not hesitate to reach out to our office.')
    c_run.font.size = Pt(11)

    doc.add_paragraph('')
    sincerely = doc.add_paragraph()
    s_run = sincerely.add_run('Sincerely,')
    s_run.font.size = Pt(11)

    doc.add_paragraph('')
    doc.add_paragraph('')
    sig = doc.add_paragraph()
    sig_run = sig.add_run('Victoria Morrison, Esq.')
    sig_run.font.size = Pt(11)
    sig_run.bold = True

    title_line = doc.add_paragraph()
    t_run = title_line.add_run('Managing Partner')
    t_run.font.size = Pt(11)

    firm_line = doc.add_paragraph()
    f_run = firm_line.add_run('Morrison, Chen & Associates')
    f_run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
