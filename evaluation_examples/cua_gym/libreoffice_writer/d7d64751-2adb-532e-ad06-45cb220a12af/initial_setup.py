"""
Initial Setup: Contact page document with icons to be inserted
Task ID: writer_obj_068
Domain: libreoffice_writer

Creates:
  - /home/user/Desktop/contact_page.docx  (document with Contact Information heading and placeholder text, NO images yet)
  - /home/user/Desktop/icon_email.png     (100x100 icon)
  - /home/user/Desktop/icon_phone.png     (100x100 icon)
  - /home/user/Desktop/icon_web.png       (100x100 icon)
  - /home/user/Desktop/icon_location.png  (100x100 icon)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_068'
OUTPUT = f'{WORKDIR}/contact_page.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_icon(path, color, label, size=(100, 100)):
    """Create a simple colored icon PNG with a label."""
    img = Image.new('RGBA', size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    # Draw filled circle as icon background
    margin = 5
    draw.ellipse(
        [margin, margin, size[0] - margin, size[1] - margin],
        fill=color,
        outline=(80, 80, 80, 255),
    )
    # Draw a simple symbol letter in the center
    text_color = (255, 255, 255, 255)
    # Use default font for simplicity
    bbox = draw.textbbox((0, 0), label)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    text_x = (size[0] - text_w) // 2
    text_y = (size[1] - text_h) // 2
    draw.text((text_x, text_y), label, fill=text_color)
    img.save(path, 'PNG')


def create_icons():
    """Create four 100x100 icon PNG files on the Desktop."""
    icons = [
        ('icon_email.png',    (52, 152, 219),  '@'),   # Blue - email
        ('icon_phone.png',    (39, 174, 96),   'P'),   # Green - phone
        ('icon_web.png',      (155, 89, 182),  'W'),   # Purple - web
        ('icon_location.png', (231, 76, 60),   'L'),   # Red - location
    ]
    for filename, color, label in icons:
        path = os.path.join(WORKDIR, filename)
        create_icon(path, color, label)
        print(f'  Created icon: {path}')


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    # Create icon files
    print('Creating icon PNG files...')
    create_icons()

    # Create the initial document
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Page 1 content: Contact Information heading
    heading = doc.add_heading('Contact Information', level=1)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(12)

    # Subtitle / intro text
    intro = doc.add_paragraph(
        'Please use the contact details below to reach our support team.'
    )
    intro.paragraph_format.space_after = Pt(8)

    # Placeholder paragraph where agent will insert icons
    placeholder = doc.add_paragraph(
        'Icons will appear here: [Email]   [Phone]   [Web]   [Location]'
    )
    placeholder.paragraph_format.space_before = Pt(6)
    placeholder.paragraph_format.space_after = Pt(8)

    # Additional contact info block
    doc.add_paragraph('Email: support@example.com')
    doc.add_paragraph('Phone: +1 (800) 555-0123')
    doc.add_paragraph('Website: www.example.com')
    doc.add_paragraph('Address: 123 Business Park, Suite 400, San Francisco, CA 94105')

    # Footer note
    footer_para = doc.add_paragraph(
        '\nFor urgent matters, please call during business hours: Monday-Friday, 9 AM – 5 PM PST.'
    )
    footer_para.paragraph_format.space_before = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial document created: {OUTPUT}')

    # GUI-ready startup: open document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
