"""
Reward Script: Insert four inline icons with correct size, alt text, and tab separation
Task ID: writer_obj_068
Domain: libreoffice_writer
Scoring:
  Component 1: 4 inline images present in document (0.35 pts)
  Component 2: Each image is 1.5cm x 1.5cm (540000 x 540000 EMU) (0.25 pts)
  Component 3: Alt text matches function labels for all 4 images (0.25 pts)
  Component 4: Tab characters separate images in the icon paragraph (0.15 pts)
  Total: 1.0
"""

import os
from docx import Document
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_068'
FILE_PATH = f'{WORKDIR}/contact_page.docx'

# Expected alt text values for the four icons (order-independent check)
EXPECTED_ALT_TEXTS = {'Email icon', 'Phone icon', 'Web icon', 'Location icon'}

# Expected image dimension in EMU: 1.5cm = 540000 EMU
# (1.5 cm / 2.54 cm per inch) * 914400 EMU per inch = 540000 EMU
EXPECTED_SIZE_EMU = 540000
# Allow a small tolerance for rounding (±5000 EMU ~ 0.05mm)
SIZE_TOLERANCE = 5000


def get_inline_images_in_doc(doc):
    """
    Extract all inline images from the document body.
    Returns a list of dicts with keys: cx, cy, alt_text, in_para_index
    """
    NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    NS_W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    images = []
    for para_idx, para in enumerate(doc.paragraphs):
        para_xml = para._element
        # Find all inline drawing elements
        for drawing in para_xml.iter(f'{{{NS_W}}}drawing'):
            for inline in drawing.iter(f'{{{NS_WP}}}inline'):
                # Get extent
                extent = inline.find(f'{{{NS_WP}}}extent')
                cx = int(extent.get('cx')) if extent is not None else None
                cy = int(extent.get('cy')) if extent is not None else None
                # Get alt text (docPr descr attribute)
                docPr = inline.find(f'{{{NS_WP}}}docPr')
                alt_text = docPr.get('descr', '') if docPr is not None else ''
                images.append({
                    'cx': cx,
                    'cy': cy,
                    'alt_text': alt_text,
                    'para_index': para_idx
                })
    return images


def count_tabs_in_icon_paragraph(doc, para_index):
    """
    Count the number of tab characters in the given paragraph.
    Tabs in .docx are represented as <w:tab/> elements within runs.
    """
    NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    para = doc.paragraphs[para_index]
    tab_count = 0
    for run_el in para._element.iter(f'{{{NS_W}}}r'):
        for tab_el in run_el.iter(f'{{{NS_W}}}tab'):
            tab_count += 1
    return tab_count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: file exists and has paragraphs
    if not doc.paragraphs:
        print("CRITICAL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: 4 inline images exist in the document (0.35 pts) ---
    # This FAILS on initial (no images) and PASSES on golden (4 images)
    try:
        images = get_inline_images_in_doc(doc)
        image_count = len(images)
        if image_count == 4:
            print(f"PASS: Component 1 — 4 inline images found (0.35 pts)")
            total_score += 0.35
        elif image_count > 0:
            # Partial: some images present but not all 4
            partial = round(0.35 * image_count / 4, 3)
            print(f"PARTIAL: Component 1 — {image_count}/4 images found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — expected 4 inline images, found {image_count}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")
        images = []

    # --- Component 2: Each image is 1.5cm x 1.5cm = 540000 x 540000 EMU (0.25 pts) ---
    # This FAILS on initial (no images) and PASSES on golden (all images sized correctly)
    try:
        if images:
            correct_size_count = 0
            for img in images:
                cx = img['cx']
                cy = img['cy']
                if (cx is not None and cy is not None and
                        abs(cx - EXPECTED_SIZE_EMU) <= SIZE_TOLERANCE and
                        abs(cy - EXPECTED_SIZE_EMU) <= SIZE_TOLERANCE):
                    correct_size_count += 1
                else:
                    print(f"FAIL: Component 2 — image '{img['alt_text']}' has size {cx}x{cy} EMU, expected ~{EXPECTED_SIZE_EMU}x{EXPECTED_SIZE_EMU}")
            if correct_size_count == len(images) and len(images) == 4:
                print(f"PASS: Component 2 — all 4 images are 1.5cm x 1.5cm ({EXPECTED_SIZE_EMU} EMU) (0.25 pts)")
                total_score += 0.25
            elif correct_size_count > 0:
                partial = round(0.25 * correct_size_count / 4, 3)
                print(f"PARTIAL: Component 2 — {correct_size_count}/4 images have correct size ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — no images have correct 1.5cm x 1.5cm size")
        else:
            print("FAIL: Component 2 — skipped (no images found)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: Alt text matches function labels for all 4 images (0.25 pts) ---
    # Expected: 'Email icon', 'Phone icon', 'Web icon', 'Location icon'
    # This FAILS on initial (no images, no alt text) and PASSES on golden
    try:
        if images:
            found_alt_texts = {img['alt_text'].strip() for img in images if img['alt_text']}
            matching = found_alt_texts & EXPECTED_ALT_TEXTS
            match_count = len(matching)
            missing = EXPECTED_ALT_TEXTS - found_alt_texts
            if match_count == 4:
                print(f"PASS: Component 3 — all 4 alt texts correct: {sorted(found_alt_texts)} (0.25 pts)")
                total_score += 0.25
            elif match_count > 0:
                partial = round(0.25 * match_count / 4, 3)
                print(f"PARTIAL: Component 3 — {match_count}/4 alt texts correct, missing: {sorted(missing)} ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — no matching alt texts found. Found: {sorted(found_alt_texts)}, expected: {sorted(EXPECTED_ALT_TEXTS)}")
        else:
            print("FAIL: Component 3 — skipped (no images found)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # --- Component 4: Tab characters separate images in the icon paragraph (0.15 pts) ---
    # There should be at least 3 tab characters between 4 images.
    # This FAILS on initial (no image paragraph with tabs) and PASSES on golden.
    try:
        if images:
            # Find the paragraph that contains the images
            icon_para_indices = list({img['para_index'] for img in images})
            tab_count = 0
            for pidx in icon_para_indices:
                tab_count += count_tabs_in_icon_paragraph(doc, pidx)
            if tab_count >= 3:
                print(f"PASS: Component 4 — {tab_count} tab characters found between images (need >= 3) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — expected >= 3 tab characters between images, found {tab_count}")
        else:
            print("FAIL: Component 4 — skipped (no images found)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 3), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path on the VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
