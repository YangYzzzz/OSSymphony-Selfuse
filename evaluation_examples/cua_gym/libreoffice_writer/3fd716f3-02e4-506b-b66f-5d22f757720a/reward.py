"""
Reward Script: Add blank lines after each sentence in the 12-sentence clause paragraph
Task ID: osworld_writer_blank_line_insertion_010
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.5 pts): Clause section sentences split into 12 separate paragraphs
  - Component 2 (0.4 pts): Each sentence paragraph followed by an empty paragraph (12 blanks)
  - Component 3 (0.1 pts): Text content of the 12 sentences is preserved

The initial document has all 12 sentences in ONE single paragraph (Para 10).
The golden document has each sentence as its own paragraph, with an empty paragraph
after each, resulting in 24 paragraphs replacing the original single one.
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_blank_line_insertion_010'

# First sentence keywords used for content integrity verification
SENTENCE_PREFIXES = [
    "The Service Provider shall not be liable",
    "Client acknowledges that Service Provider",
    "Service Provider shall indemnify and hold harmless Client",
    "Client shall indemnify and hold harmless Service Provider",
    "Each party shall promptly notify the other",
    "The indemnifying party shall have the right",
    "The indemnified party shall cooperate fully",
    "Neither party shall settle any claim",
    "The limitations of liability set forth herein",
    "These limitations shall survive the termination",
    "Service Provider maintains professional liability insurance",
    "Client shall be named as an additional insured",
]


def find_section4_heading_index(doc):
    """Find the index of the 'Section 4. Liability and Indemnification Clause' heading."""
    for i, para in enumerate(doc.paragraphs):
        if "Section 4" in para.text and "Liability" in para.text:
            return i
    return None


def find_section5_heading_index(doc):
    """Find the index of the 'Section 5. Term and Termination' heading (end of clause section)."""
    for i, para in enumerate(doc.paragraphs):
        if "Section 5" in para.text and "Termination" in para.text:
            return i
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate Section 4 heading and Section 5 heading to bound the clause section
    sec4_idx = find_section4_heading_index(doc)
    sec5_idx = find_section5_heading_index(doc)

    if sec4_idx is None or sec5_idx is None:
        print(f"CRITICAL: Could not find Section 4 or Section 5 heading in document")
        print(f"  sec4_idx={sec4_idx}, sec5_idx={sec5_idx}")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Section 4 heading at para index {sec4_idx}, Section 5 heading at para index {sec5_idx}")

    # Extract paragraphs that belong to the clause section (between the two headings)
    clause_paras = doc.paragraphs[sec4_idx + 1 : sec5_idx]
    print(f"INFO: Clause section contains {len(clause_paras)} paragraphs total (between headings)")

    # Classify paragraphs
    sentence_paras = [p for p in clause_paras if p.text.strip()]
    empty_paras = [p for p in clause_paras if not p.text.strip()]

    print(f"INFO: Non-empty paragraphs: {len(sentence_paras)}, Empty paragraphs: {len(empty_paras)}")

    # Component 1: Clause section sentences split into 12 separate paragraphs (0.5 points)
    # In initial state: 1 paragraph contains all 12 sentences.
    # In golden state: 12 separate non-empty paragraphs.
    try:
        num_sentence_paras = len(sentence_paras)
        if num_sentence_paras == 12:
            print(f"PASS: Component 1 — exactly 12 sentence paragraphs found in clause section (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — expected 12 sentence paragraphs, found {num_sentence_paras}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Each sentence paragraph followed by an empty paragraph (0.4 points)
    # In initial state: 0 blank lines (all sentences in one paragraph).
    # In golden state: each sentence para is immediately followed by an empty para = 12 empty paras.
    try:
        # Verify alternating pattern: sentence, empty, sentence, empty, ...
        # The clause section should be: [sent, empty, sent, empty, ..., sent, empty] = 24 paras
        expected_total = 24  # 12 sentences + 12 empty lines

        if len(clause_paras) != expected_total:
            print(f"FAIL: Component 2 — expected {expected_total} total clause paragraphs "
                  f"(12 sentences + 12 blanks), found {len(clause_paras)}")
        else:
            # Verify the alternating pattern: even indices (0,2,4,...) are sentence, odd are empty
            violations = []
            for j in range(0, 24, 2):
                sentence_para = clause_paras[j]
                blank_para = clause_paras[j + 1]
                if not sentence_para.text.strip():
                    violations.append(f"  Para {j} (expected sentence) is empty")
                if blank_para.text.strip():
                    violations.append(f"  Para {j+1} (expected blank) is not empty: {blank_para.text.strip()[:50]!r}")

            if len(violations) == 0:
                print("PASS: Component 2 — alternating pattern verified: "
                      "each sentence followed by blank line (0.4 pts)")
                total_score += 0.4
            else:
                print("FAIL: Component 2 — alternating pattern not correct:")
                for v in violations:
                    print(v)
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text content of the 12 sentences is preserved (0.1 points)
    # The text of each sentence should begin with the expected prefix.
    try:
        sentences_found = 0
        for prefix in SENTENCE_PREFIXES:
            found = any(p.text.strip().startswith(prefix) for p in sentence_paras)
            if found:
                sentences_found += 1
            else:
                print(f"  NOTE: Sentence starting with '{prefix[:50]}...' not found as individual paragraph")

        if sentences_found == 12:
            print(f"PASS: Component 3 — all 12 sentences present and content intact (0.1 pts)")
            total_score += 0.1
        elif sentences_found >= 10:
            print(f"PARTIAL Component 3 — {sentences_found}/12 sentence texts found ({0.05:.2f} pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — only {sentences_found}/12 sentence texts found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
