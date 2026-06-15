"""
Initial Setup: Legal contract with 12-sentence clause paragraph (no blank lines between sentences)
Task ID: osworld_writer_blank_line_insertion_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_blank_line_insertion_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Document Title ---
    title = doc.add_heading('SERVICE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    doc.add_paragraph(
        'This Service Agreement ("Agreement") is entered into as of March 1, 2025, '
        'by and between Hartwell Consulting Group, LLC, a Delaware limited liability company '
        '("Service Provider"), and Meridian Financial Partners, Inc., a California corporation ("Client").'
    )

    doc.add_paragraph(
        'WHEREAS, Service Provider desires to provide certain professional consulting services to Client; '
        'and WHEREAS, Client desires to retain the services of Service Provider on the terms and conditions '
        'set forth herein; NOW, THEREFORE, in consideration of the mutual covenants contained herein, '
        'the parties agree as follows:'
    )

    # --- Section 1: Definitions ---
    doc.add_heading('Section 1. Definitions', level=1)
    doc.add_paragraph(
        '"Services" means the professional consulting, advisory, and related services described in Exhibit A '
        'attached hereto and incorporated herein by reference. "Deliverables" means all work product, '
        'reports, analyses, and documents produced by Service Provider under this Agreement. '
        '"Confidential Information" means all non-public information disclosed by either party to the other.'
    )

    # --- Section 2: Scope of Services ---
    doc.add_heading('Section 2. Scope of Services', level=1)
    doc.add_paragraph(
        'Service Provider shall perform the Services described in Exhibit A with reasonable care and skill. '
        'Service Provider may engage qualified subcontractors to assist with certain portions of the Services, '
        'provided that Service Provider shall remain responsible for the performance of all Services.'
    )

    # --- Section 3: Compensation ---
    doc.add_heading('Section 3. Compensation and Payment Terms', level=1)
    doc.add_paragraph(
        'Client shall pay Service Provider a monthly retainer fee of $12,500.00 for the duration of this Agreement. '
        'Invoices shall be issued on the first business day of each calendar month. '
        'Payment shall be due within thirty (30) days of the invoice date. '
        'Late payments shall accrue interest at the rate of 1.5% per month.'
    )

    # --- Section 4: Clause (THE KEY SECTION — 12 sentences in ONE paragraph) ---
    doc.add_heading('Section 4. Liability and Indemnification Clause', level=1)

    # All 12 sentences in a SINGLE paragraph (no blank lines between them — that is the task)
    clause_sentences = [
        'The Service Provider shall not be liable for any indirect, incidental, special, or consequential damages arising out of or related to this Agreement.',
        'Client acknowledges that Service Provider\'s total cumulative liability under this Agreement shall not exceed the total fees paid by Client in the three (3) months preceding the claim.',
        'Service Provider shall indemnify and hold harmless Client from any third-party claims arising directly from Service Provider\'s gross negligence or willful misconduct.',
        'Client shall indemnify and hold harmless Service Provider from any claims arising out of Client\'s use or misuse of the Deliverables.',
        'Each party shall promptly notify the other in writing upon becoming aware of any claim or potential claim covered by this indemnification provision.',
        'The indemnifying party shall have the right to assume control of the defense of any such claim, at its own cost and expense.',
        'The indemnified party shall cooperate fully with the indemnifying party in the defense of any such claim, including providing access to relevant documents and witnesses.',
        'Neither party shall settle any claim subject to indemnification without the prior written consent of the other party, which shall not be unreasonably withheld.',
        'The limitations of liability set forth herein shall apply regardless of the form of action, whether in contract, tort, strict liability, or otherwise.',
        'These limitations shall survive the termination or expiration of this Agreement and shall continue to bind the parties indefinitely.',
        'Service Provider maintains professional liability insurance in amounts no less than $2,000,000 per occurrence throughout the term of this Agreement.',
        'Client shall be named as an additional insured on Service Provider\'s general liability policy upon written request and at no additional cost.',
    ]

    # Join all sentences into one single paragraph
    clause_text = ' '.join(clause_sentences)
    clause_para = doc.add_paragraph(clause_text)
    clause_para.paragraph_format.space_after = Pt(0)

    # --- Section 5: Term and Termination ---
    doc.add_heading('Section 5. Term and Termination', level=1)
    doc.add_paragraph(
        'This Agreement shall commence on March 1, 2025, and shall continue for a period of one (1) year '
        'unless earlier terminated in accordance with this Section. Either party may terminate this Agreement '
        'upon sixty (60) days\' prior written notice to the other party. Upon termination, Client shall pay '
        'Service Provider for all Services rendered through the effective date of termination.'
    )

    # --- Section 6: Confidentiality ---
    doc.add_heading('Section 6. Confidentiality', level=1)
    doc.add_paragraph(
        'Each party agrees to maintain the confidentiality of the other party\'s Confidential Information '
        'and to use such information solely for the purposes of this Agreement. The receiving party shall '
        'protect Confidential Information with at least the same degree of care used to protect its own '
        'confidential information, but in no event less than reasonable care. The obligations of confidentiality '
        'shall survive termination of this Agreement for a period of five (5) years.'
    )

    # --- Signature Block ---
    doc.add_heading('Signatures', level=1)
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    doc.add_paragraph('')
    sig_para = doc.add_paragraph('Hartwell Consulting Group, LLC\t\t\tMeridian Financial Partners, Inc.')
    sig_para.paragraph_format.space_before = Pt(24)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
