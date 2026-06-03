"""
Reward Script: Merge two documents by comparing with Alice's and Bob's edits
Task ID: writer_fp_016
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Alice's correction to Executive Overview paragraph (added enterprise clients text)
  Component 2 (0.25): Alice's correction to Operations paragraph (2.9 days + routing algorithm)
  Component 3 (0.25): Bob's new section 4.1 Diversity and Inclusion (heading + content)
  Component 4 (0.25): Bob's table reformatting (extra column + new D&I table)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_016'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task merges Alice's and Bob's edits into the base document.
    Alice's edits: corrections to paragraphs about Executive Overview and Operations.
    Bob's edits: new section 4.1 (D&I) and reformatted tables (extra columns + new table).
    All tracked changes should be accepted in the final document.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = [p.text for p in doc.paragraphs]
    full_text = '\n'.join(all_text)

    # Component 1: Alice's correction to Executive Overview (0.25 points)
    # In initial: paragraph about revenue ends with "...Southeast Asian corridor."
    # In golden: adds "...and strong demand from enterprise clients in Japan and South Korea."
    try:
        alice_correction_1_found = False
        for p in doc.paragraphs:
            if "enterprise clients in Japan and South Korea" in p.text:
                alice_correction_1_found = True
                break

        if alice_correction_1_found:
            print(f"PASS: Component 1 - Alice's correction found: enterprise clients in Japan and South Korea (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - Alice's correction missing: 'enterprise clients in Japan and South Korea' not found")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Alice's correction to Operations paragraph (0.25 points)
    # In initial: "3.2 days, slightly above the target"
    # In golden: "2.9 days, meeting the target" AND added "new routing algorithm deployed in November"
    try:
        has_corrected_days = False
        has_routing_algorithm = False
        for p in doc.paragraphs:
            if "2.9 days" in p.text and "meeting the target" in p.text:
                has_corrected_days = True
            if "routing algorithm" in p.text and "November" in p.text:
                has_routing_algorithm = True

        # Must NOT have the old incorrect text
        has_old_text = "3.2 days" in full_text and "slightly above" in full_text

        if has_corrected_days and has_routing_algorithm and not has_old_text:
            print(f"PASS: Component 2 - Alice's Operations correction: 2.9 days + routing algorithm (0.25 pts)")
            total_score += 0.25
        elif has_corrected_days or has_routing_algorithm:
            partial = 0.125
            print(f"PARTIAL: Component 2 - Only partial correction found (corrected_days={has_corrected_days}, routing={has_routing_algorithm}, old_removed={not has_old_text}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - Alice's Operations correction missing")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Bob's new section 4.1 Diversity and Inclusion (0.25 points)
    # Golden has a Heading 2 "4.1 Diversity and Inclusion Initiatives" plus content paragraphs
    try:
        has_section_41_heading = False
        has_di_content = False

        for p in doc.paragraphs:
            if "4.1" in p.text and "Diversity" in p.text and "Inclusion" in p.text:
                has_section_41_heading = True
            if "mentorship program" in p.text.lower() or "underrepresented groups" in p.text.lower():
                has_di_content = True
            if "employee resource groups" in p.text.lower() or "ERGs" in p.text:
                has_di_content = True

        if has_section_41_heading and has_di_content:
            print(f"PASS: Component 3 - Bob's section 4.1 D&I heading + content (0.25 pts)")
            total_score += 0.25
        elif has_section_41_heading:
            print(f"PARTIAL: Component 3 - Section 4.1 heading found but content missing (0.125 pts)")
            total_score += 0.125
        else:
            print(f"FAIL: Component 3 - Bob's section 4.1 not found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Bob's table reformatting (0.25 points)
    # Initial table 0: 4 columns (Metric, Q3, Q4, Change)
    # Golden table 0: 5 columns (Metric, Q3, Q4, QoQ Change, YoY Change)
    # Golden also adds a new Table 1 for D&I metrics
    try:
        tables = doc.tables
        has_expanded_table = False
        has_di_table = False

        if len(tables) >= 1:
            t0 = tables[0]
            # Check for 5 columns and YoY Change header
            if len(t0.columns) >= 5:
                header_cells = [c.text.strip() for c in t0.rows[0].cells]
                if "YoY Change" in header_cells:
                    has_expanded_table = True

        if len(tables) >= 2:
            t1 = tables[1]
            # Check the D&I table has expected content
            all_table_text = ' '.join(c.text for row in t1.rows for c in row.cells)
            if "Initiative" in all_table_text or "Mentorship" in all_table_text or "ERG" in all_table_text:
                has_di_table = True

        if has_expanded_table and has_di_table:
            print(f"PASS: Component 4 - Bob's table reformatting: 5-col table + D&I table (0.25 pts)")
            total_score += 0.25
        elif has_expanded_table or has_di_table:
            partial = 0.125
            print(f"PARTIAL: Component 4 - Only partial table changes (expanded={has_expanded_table}, di_table={has_di_table}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - Bob's table reformatting not found (tables={len(tables)}, cols in t0={len(tables[0].columns) if tables else 'N/A'})")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
