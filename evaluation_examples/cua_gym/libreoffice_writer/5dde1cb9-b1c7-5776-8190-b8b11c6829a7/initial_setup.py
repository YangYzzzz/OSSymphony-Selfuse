"""
Initial Setup: Merge two reviewer documents into base report
Task ID: writer_fp_016
Domain: libreoffice_writer

Creates:
  /home/user/writer_fp_016.docx - Base report (original, no reviewer changes)
  /home/user/Documents/report_edits_alice.docx - Alice's edits (paragraphs 3,7 corrected)
  /home/user/Documents/report_edits_bob.docx - Bob's edits (new section 4.1, reformatted tables)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_016'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DOCS_DIR = f'{WORKDIR}/Documents'
ALICE_FILE = f'{DOCS_DIR}/report_edits_alice.docx'
BOB_FILE = f'{DOCS_DIR}/report_edits_bob.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body_paragraph(doc, text, bold_first_sentence=False):
    """Add a body paragraph with realistic text."""
    para = doc.add_paragraph()
    if bold_first_sentence and '.' in text:
        idx = text.index('.') + 1
        run_bold = para.add_run(text[:idx])
        run_bold.bold = True
        run_normal = para.add_run(text[idx:])
    else:
        para.add_run(text)
    return para


def create_base_report():
    """Create the original base report document."""
    doc = Document()

    # Title
    title = doc.add_heading('Quarterly Performance Report - Q4 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 1
    add_body_paragraph(doc,
        'This report presents a comprehensive analysis of the organizational performance '
        'metrics for the fourth quarter of fiscal year 2025. The data collected spans all '
        'operational divisions and reflects activities from October 1 through December 31, 2025.')

    # Paragraph 2
    add_body_paragraph(doc,
        'Executive leadership has reviewed the preliminary findings and endorsed the '
        'methodology used for data collection. The sampling framework covers 94% of all '
        'active business units across North America, Europe, and Asia-Pacific regions.')

    # Section 1: Overview
    add_heading_styled(doc, '1. Executive Overview', level=1)

    # Paragraph 3 (Alice will correct this)
    add_body_paragraph(doc,
        'Total revenue for Q4 reached $12.8 million, representing a 7% increase over '
        'the previous quarter. Operating expenses remained stable at $9.1 million, yielding '
        'an operating margin of approximately 28.9%. The primary growth drivers were the '
        'expanded product line in consumer electronics and increased market penetration in '
        'the Southeast Asian corridor.')

    # Paragraph 4
    add_body_paragraph(doc,
        'Customer acquisition costs decreased by 12% compared to Q3, largely attributed '
        'to the digital marketing optimization initiative launched in September. The average '
        'customer lifetime value increased to $2,340, up from $2,180 in the prior quarter.')

    # Section 2: Financial Performance
    add_heading_styled(doc, '2. Financial Performance', level=1)

    # Paragraph 5
    add_body_paragraph(doc,
        'The finance division reported strong cash flow generation during Q4. Net cash from '
        'operations totaled $3.2 million, while capital expenditures were limited to $890,000 '
        'for infrastructure upgrades and technology investments.')

    # Financial Table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Metric', 'Q3 2025', 'Q4 2025', 'Change']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ['Revenue', '$11.96M', '$12.80M', '+7.0%'],
        ['Operating Expenses', '$9.05M', '$9.10M', '+0.6%'],
        ['Net Income', '$2.18M', '$2.77M', '+27.1%'],
        ['EBITDA', '$3.45M', '$3.92M', '+13.6%'],
        ['Free Cash Flow', '$2.10M', '$2.31M', '+10.0%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Paragraph 6
    add_body_paragraph(doc,
        'The improvement in net income was primarily driven by higher gross margins in the '
        'software licensing segment and reduced overhead costs following the restructuring '
        'completed in Q2 2025.')

    # Section 3: Operations
    add_heading_styled(doc, '3. Operations and Supply Chain', level=1)

    # Paragraph 7 (Alice will correct this)
    add_body_paragraph(doc,
        'Supply chain operations maintained efficiency throughout Q4 despite global logistics '
        'challenges. The average order fulfillment time was 3.2 days, slightly above the '
        'target of 3.0 days. Inventory turnover ratio stood at 6.8, reflecting healthy stock '
        'management practices across all distribution centers.')

    # Paragraph 8
    add_body_paragraph(doc,
        'The warehouse automation project in the Portland facility reached 85% completion by '
        'end of quarter. Early metrics show a 22% improvement in picking accuracy and 15% '
        'reduction in labor costs per unit shipped compared to the manual baseline.')

    # Section 4: Human Resources (Bob will add 4.1 under this)
    add_heading_styled(doc, '4. Human Resources', level=1)

    # Paragraph 9
    add_body_paragraph(doc,
        'The workforce expanded by 45 positions during Q4, with the majority of new hires '
        'joining the engineering and customer success teams. Employee retention rate remained '
        'strong at 92%, exceeding the industry average of 87%.')

    # Paragraph 10
    add_body_paragraph(doc,
        'Training and development expenditures totaled $340,000, with 78% of staff completing '
        'at least one professional development course during the quarter. Leadership identified '
        'skill gaps in cloud architecture and data analytics as priorities for Q1 2026.')

    # Section 5: Outlook
    add_heading_styled(doc, '5. Outlook and Recommendations', level=1)

    # Paragraph 11
    add_body_paragraph(doc,
        'Based on current trends and pipeline analysis, management projects Q1 2026 revenue '
        'of $13.5 million, representing continued sequential growth. Key risks include '
        'potential tariff changes in the Asia-Pacific region and increasing competition in '
        'the mid-market enterprise software segment.')

    # Paragraph 12
    add_body_paragraph(doc,
        'The strategic planning committee recommends accelerating investment in artificial '
        'intelligence capabilities, expanding the partner ecosystem, and completing the '
        'warehouse automation rollout by end of Q2 2026.')

    doc.save(OUTPUT)
    print(f'Base report created: {OUTPUT}')


def create_alice_edits():
    """Create Alice's version with corrections to paragraphs 3 and 7."""
    doc = Document()

    # Title
    title = doc.add_heading('Quarterly Performance Report - Q4 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 1 (unchanged)
    add_body_paragraph(doc,
        'This report presents a comprehensive analysis of the organizational performance '
        'metrics for the fourth quarter of fiscal year 2025. The data collected spans all '
        'operational divisions and reflects activities from October 1 through December 31, 2025.')

    # Paragraph 2 (unchanged)
    add_body_paragraph(doc,
        'Executive leadership has reviewed the preliminary findings and endorsed the '
        'methodology used for data collection. The sampling framework covers 94% of all '
        'active business units across North America, Europe, and Asia-Pacific regions.')

    # Section 1
    add_heading_styled(doc, '1. Executive Overview', level=1)

    # Paragraph 3 - ALICE'S CORRECTION: Fixed revenue figure and margin calculation
    add_body_paragraph(doc,
        'Total revenue for Q4 reached $12.8 million, representing a 7% increase over '
        'the previous quarter. Operating expenses remained stable at $9.1 million, yielding '
        'an operating margin of approximately 28.9%. The primary growth drivers were the '
        'expanded product line in consumer electronics, increased market penetration in '
        'the Southeast Asian corridor, and strong demand from enterprise clients in Japan and South Korea.')

    # Paragraph 4 (unchanged)
    add_body_paragraph(doc,
        'Customer acquisition costs decreased by 12% compared to Q3, largely attributed '
        'to the digital marketing optimization initiative launched in September. The average '
        'customer lifetime value increased to $2,340, up from $2,180 in the prior quarter.')

    # Section 2
    add_heading_styled(doc, '2. Financial Performance', level=1)

    # Paragraph 5 (unchanged)
    add_body_paragraph(doc,
        'The finance division reported strong cash flow generation during Q4. Net cash from '
        'operations totaled $3.2 million, while capital expenditures were limited to $890,000 '
        'for infrastructure upgrades and technology investments.')

    # Financial Table (unchanged)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Metric', 'Q3 2025', 'Q4 2025', 'Change']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ['Revenue', '$11.96M', '$12.80M', '+7.0%'],
        ['Operating Expenses', '$9.05M', '$9.10M', '+0.6%'],
        ['Net Income', '$2.18M', '$2.77M', '+27.1%'],
        ['EBITDA', '$3.45M', '$3.92M', '+13.6%'],
        ['Free Cash Flow', '$2.10M', '$2.31M', '+10.0%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Paragraph 6 (unchanged)
    add_body_paragraph(doc,
        'The improvement in net income was primarily driven by higher gross margins in the '
        'software licensing segment and reduced overhead costs following the restructuring '
        'completed in Q2 2025.')

    # Section 3
    add_heading_styled(doc, '3. Operations and Supply Chain', level=1)

    # Paragraph 7 - ALICE'S CORRECTION: Updated fulfillment time and added detail
    add_body_paragraph(doc,
        'Supply chain operations maintained efficiency throughout Q4 despite global logistics '
        'challenges. The average order fulfillment time was 2.9 days, meeting the '
        'target of 3.0 days. Inventory turnover ratio stood at 6.8, reflecting healthy stock '
        'management practices across all distribution centers. The improvement was attributed '
        'to the new routing algorithm deployed in November.')

    # Paragraph 8 (unchanged)
    add_body_paragraph(doc,
        'The warehouse automation project in the Portland facility reached 85% completion by '
        'end of quarter. Early metrics show a 22% improvement in picking accuracy and 15% '
        'reduction in labor costs per unit shipped compared to the manual baseline.')

    # Section 4
    add_heading_styled(doc, '4. Human Resources', level=1)

    # Paragraph 9 (unchanged)
    add_body_paragraph(doc,
        'The workforce expanded by 45 positions during Q4, with the majority of new hires '
        'joining the engineering and customer success teams. Employee retention rate remained '
        'strong at 92%, exceeding the industry average of 87%.')

    # Paragraph 10 (unchanged)
    add_body_paragraph(doc,
        'Training and development expenditures totaled $340,000, with 78% of staff completing '
        'at least one professional development course during the quarter. Leadership identified '
        'skill gaps in cloud architecture and data analytics as priorities for Q1 2026.')

    # Section 5
    add_heading_styled(doc, '5. Outlook and Recommendations', level=1)

    # Paragraph 11 (unchanged)
    add_body_paragraph(doc,
        'Based on current trends and pipeline analysis, management projects Q1 2026 revenue '
        'of $13.5 million, representing continued sequential growth. Key risks include '
        'potential tariff changes in the Asia-Pacific region and increasing competition in '
        'the mid-market enterprise software segment.')

    # Paragraph 12 (unchanged)
    add_body_paragraph(doc,
        'The strategic planning committee recommends accelerating investment in artificial '
        'intelligence capabilities, expanding the partner ecosystem, and completing the '
        'warehouse automation rollout by end of Q2 2026.')

    os.makedirs(os.path.dirname(ALICE_FILE), exist_ok=True)
    doc.save(ALICE_FILE)
    print(f'Alice edits created: {ALICE_FILE}')


def create_bob_edits():
    """Create Bob's version with new section 4.1 and reformatted tables."""
    doc = Document()

    # Title
    title = doc.add_heading('Quarterly Performance Report - Q4 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 1 (unchanged)
    add_body_paragraph(doc,
        'This report presents a comprehensive analysis of the organizational performance '
        'metrics for the fourth quarter of fiscal year 2025. The data collected spans all '
        'operational divisions and reflects activities from October 1 through December 31, 2025.')

    # Paragraph 2 (unchanged)
    add_body_paragraph(doc,
        'Executive leadership has reviewed the preliminary findings and endorsed the '
        'methodology used for data collection. The sampling framework covers 94% of all '
        'active business units across North America, Europe, and Asia-Pacific regions.')

    # Section 1
    add_heading_styled(doc, '1. Executive Overview', level=1)

    # Paragraph 3 (unchanged from base)
    add_body_paragraph(doc,
        'Total revenue for Q4 reached $12.8 million, representing a 7% increase over '
        'the previous quarter. Operating expenses remained stable at $9.1 million, yielding '
        'an operating margin of approximately 28.9%. The primary growth drivers were the '
        'expanded product line in consumer electronics and increased market penetration in '
        'the Southeast Asian corridor.')

    # Paragraph 4 (unchanged)
    add_body_paragraph(doc,
        'Customer acquisition costs decreased by 12% compared to Q3, largely attributed '
        'to the digital marketing optimization initiative launched in September. The average '
        'customer lifetime value increased to $2,340, up from $2,180 in the prior quarter.')

    # Section 2
    add_heading_styled(doc, '2. Financial Performance', level=1)

    # Paragraph 5 (unchanged)
    add_body_paragraph(doc,
        'The finance division reported strong cash flow generation during Q4. Net cash from '
        'operations totaled $3.2 million, while capital expenditures were limited to $890,000 '
        'for infrastructure upgrades and technology investments.')

    # BOB'S REFORMATTED TABLE: Added header formatting and YoY column
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    headers = ['Metric', 'Q3 2025', 'Q4 2025', 'QoQ Change', 'YoY Change']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(11)
    data = [
        ['Revenue', '$11.96M', '$12.80M', '+7.0%', '+18.5%'],
        ['Operating Expenses', '$9.05M', '$9.10M', '+0.6%', '+3.2%'],
        ['Net Income', '$2.18M', '$2.77M', '+27.1%', '+41.3%'],
        ['EBITDA', '$3.45M', '$3.92M', '+13.6%', '+25.8%'],
        ['Free Cash Flow', '$2.10M', '$2.31M', '+10.0%', '+19.4%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Paragraph 6 (unchanged)
    add_body_paragraph(doc,
        'The improvement in net income was primarily driven by higher gross margins in the '
        'software licensing segment and reduced overhead costs following the restructuring '
        'completed in Q2 2025.')

    # Section 3
    add_heading_styled(doc, '3. Operations and Supply Chain', level=1)

    # Paragraph 7 (unchanged from base)
    add_body_paragraph(doc,
        'Supply chain operations maintained efficiency throughout Q4 despite global logistics '
        'challenges. The average order fulfillment time was 3.2 days, slightly above the '
        'target of 3.0 days. Inventory turnover ratio stood at 6.8, reflecting healthy stock '
        'management practices across all distribution centers.')

    # Paragraph 8 (unchanged)
    add_body_paragraph(doc,
        'The warehouse automation project in the Portland facility reached 85% completion by '
        'end of quarter. Early metrics show a 22% improvement in picking accuracy and 15% '
        'reduction in labor costs per unit shipped compared to the manual baseline.')

    # Section 4
    add_heading_styled(doc, '4. Human Resources', level=1)

    # BOB'S NEW SECTION 4.1
    add_heading_styled(doc, '4.1 Diversity and Inclusion Initiatives', level=2)

    add_body_paragraph(doc,
        'The diversity and inclusion program launched in Q3 continued to gain momentum during Q4. '
        'Representation of underrepresented groups in technical roles increased from 31% to 35%, '
        'surpassing the annual target of 33%. The mentorship program paired 62 junior employees '
        'with senior leaders across all departments.')

    add_body_paragraph(doc,
        'Employee resource groups (ERGs) hosted 14 events during Q4, with average attendance of '
        '45 participants per event. The annual inclusion survey showed a satisfaction score of '
        '4.2 out of 5.0, up from 3.8 in the previous year. Management has allocated an additional '
        '$150,000 for D&I initiatives in the Q1 2026 budget.')

    # BOB'S REFORMATTED D&I TABLE
    di_table = doc.add_table(rows=5, cols=3)
    di_table.style = 'Table Grid'
    di_headers = ['Initiative', 'Q4 Result', 'Target']
    for i, h in enumerate(di_headers):
        cell = di_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(11)
    di_data = [
        ['Technical Role Representation', '35%', '33%'],
        ['Mentorship Pairs', '62', '50'],
        ['ERG Events', '14', '10'],
        ['Inclusion Survey Score', '4.2/5.0', '4.0/5.0'],
    ]
    for r, row_data in enumerate(di_data, 1):
        for c, val in enumerate(row_data):
            di_table.cell(r, c).text = val

    # Paragraph 9 (unchanged)
    add_body_paragraph(doc,
        'The workforce expanded by 45 positions during Q4, with the majority of new hires '
        'joining the engineering and customer success teams. Employee retention rate remained '
        'strong at 92%, exceeding the industry average of 87%.')

    # Paragraph 10 (unchanged)
    add_body_paragraph(doc,
        'Training and development expenditures totaled $340,000, with 78% of staff completing '
        'at least one professional development course during the quarter. Leadership identified '
        'skill gaps in cloud architecture and data analytics as priorities for Q1 2026.')

    # Section 5
    add_heading_styled(doc, '5. Outlook and Recommendations', level=1)

    # Paragraph 11 (unchanged)
    add_body_paragraph(doc,
        'Based on current trends and pipeline analysis, management projects Q1 2026 revenue '
        'of $13.5 million, representing continued sequential growth. Key risks include '
        'potential tariff changes in the Asia-Pacific region and increasing competition in '
        'the mid-market enterprise software segment.')

    # Paragraph 12 (unchanged)
    add_body_paragraph(doc,
        'The strategic planning committee recommends accelerating investment in artificial '
        'intelligence capabilities, expanding the partner ecosystem, and completing the '
        'warehouse automation rollout by end of Q2 2026.')

    os.makedirs(os.path.dirname(BOB_FILE), exist_ok=True)
    doc.save(BOB_FILE)
    print(f'Bob edits created: {BOB_FILE}')


# Create all files
create_base_report()
create_alice_edits()
create_bob_edits()

# GUI-ready startup: open the base document in LibreOffice Writer
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
