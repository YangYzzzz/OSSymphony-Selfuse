"""
Initial Setup: Remove indentation and set space_after for consulting proposal
Task ID: writer_para_079
Domain: libreoffice_writer

Creates consulting_proposal.docx with indented paragraphs (pre-task state).
Paragraphs 3-7 have first_line_indent=1.25cm; no space_after set.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'consulting_proposal'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Paragraph 1: 'Management Consulting Proposal' (Heading 1, center-aligned)
    p1 = doc.add_heading('Management Consulting Proposal', level=1)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 2: 'Prepared for: Meridian Healthcare Group' (center-aligned, normal style)
    p2 = doc.add_paragraph('Prepared for: Meridian Healthcare Group')
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraphs 3-7: Body text with first_line_indent=1.25cm (pre-task indented state)
    body_texts = [
        'Meridian Healthcare Group is seeking to optimize its operational efficiency across its network of 12 hospitals and 45 outpatient clinics.',
        'Our team of healthcare management consultants brings over 50 years of combined experience in hospital operations, revenue cycle management, and clinical workflow optimization.',
        'The proposed engagement will span 16 weeks and consist of three phases: diagnostic assessment, solution design, and implementation support.',
        'Phase 1 will involve on-site observations at four representative facilities, interviews with key stakeholders, and analysis of operational data spanning the past three fiscal years.',
        'Our fee for the complete engagement is $480,000, payable in monthly installments aligned with phase milestones.',
    ]

    for text in body_texts:
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.25)
        # No space_after set (default, task is to add 12pt)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
