"""
Reward Script: Remove indentation and set uniform space_after=12pt on body paragraphs
Task ID: writer_para_079
Domain: libreoffice_writer
Scoring:
  Component 1: first_line_indent == 0 for paragraphs 2-6 (0-indexed)  — 0.4 pts
  Component 2: left_indent explicitly == 0pt for paragraphs 2-6        — 0.3 pts
  Component 3: space_after == 12pt for paragraphs 2-6                  — 0.3 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_para_079'
FILE_PATH = f'{WORKDIR}/Desktop/consulting_proposal.docx'

# Indices (0-based) of the body paragraphs that must be reformatted
BODY_PARA_INDICES = [2, 3, 4, 5, 6]

# Tolerance for float comparison (in points)
TOLERANCE_PT = 0.5


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Remove first_line_indent from body paragraphs (originally 1.25cm / ~35.45pt),
          set left_indent = 0, and set space_after = 12pt for paragraphs 3-7 (1-indexed),
          i.e. indices 2-6 (0-indexed). Paragraphs 0 and 1 (heading and sub-heading) must
          remain untouched.
    """
    total_score = 0.0

    # Load document — if it fails, bail out immediately
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 7 paragraphs
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected at least 7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs

    # ------------------------------------------------------------------ #
    # Component 1: first_line_indent == 0 for body paragraphs (0.4 pts)
    # In the initial file paragraphs 2-6 have first_line_indent ~35.45pt.
    # After the task they must be explicitly set to 0.
    # ------------------------------------------------------------------ #
    try:
        fi_details = []
        fi_pass_count = 0
        for idx in BODY_PARA_INDICES:
            para = paragraphs[idx]
            fi = para.paragraph_format.first_line_indent
            fi_pt = fi.pt if fi is not None else None
            fi_details.append((idx, fi_pt))
            if fi_pt is not None and abs(fi_pt) <= TOLERANCE_PT:
                fi_pass_count += 1

        if fi_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 1 — first_line_indent == 0 for all body paragraphs (0.4 pts)")
            print(f"      Details: {fi_details}")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — first_line_indent NOT zero for all body paragraphs ({fi_pass_count}/{len(BODY_PARA_INDICES)} passed)")
            print(f"      Details: {fi_details}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ------------------------------------------------------------------ #
    # Component 2: left_indent explicitly set to 0pt for body paragraphs (0.3 pts)
    # In the initial file left_indent is None (not set / inherited).
    # After the task it must be explicitly set to 0.0pt (not None).
    # The golden file sets it explicitly to 0.0pt to convert to block style.
    # ------------------------------------------------------------------ #
    try:
        li_details = []
        li_pass_count = 0
        for idx in BODY_PARA_INDICES:
            para = paragraphs[idx]
            li = para.paragraph_format.left_indent
            li_pt = li.pt if li is not None else None
            li_details.append((idx, li_pt))
            # Must be explicitly set to 0.0pt (not None/unset)
            if li_pt is not None and abs(li_pt) <= TOLERANCE_PT:
                li_pass_count += 1

        if li_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 2 — left_indent explicitly == 0pt for all body paragraphs (0.3 pts)")
            print(f"      Details: {li_details}")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — left_indent NOT explicitly 0pt for all body paragraphs ({li_pass_count}/{len(BODY_PARA_INDICES)} passed)")
            print(f"      Details: {li_details}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ------------------------------------------------------------------ #
    # Component 3: space_after == 12pt for body paragraphs (0.3 pts)
    # In the initial file space_after is None (not set).
    # After the task it must be explicitly 12pt.
    # ------------------------------------------------------------------ #
    try:
        sa_details = []
        sa_pass_count = 0
        for idx in BODY_PARA_INDICES:
            para = paragraphs[idx]
            sa = para.paragraph_format.space_after
            sa_pt = sa.pt if sa is not None else None
            sa_details.append((idx, sa_pt))
            if sa_pt is not None and abs(sa_pt - 12.0) <= TOLERANCE_PT:
                sa_pass_count += 1

        if sa_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 3 — space_after == 12pt for all body paragraphs (0.3 pts)")
            print(f"      Details: {sa_details}")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — space_after NOT 12pt for all body paragraphs ({sa_pass_count}/{len(BODY_PARA_INDICES)} passed)")
            print(f"      Details: {sa_details}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against the canonical artifact path on the VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
