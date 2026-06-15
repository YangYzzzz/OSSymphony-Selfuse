"""
Initial Setup: Annual Performance Review document with basic employee info
Task ID: writer_hr_061
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_061'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('Annual Performance Review', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Fiscal Year 2025-2026')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.italic = True

    doc.add_paragraph()  # spacer

    # --- Employee Information Section ---
    info_heading = doc.add_heading('Employee Information', level=1)

    # Employee info as a simple table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'

    fields = [
        ('Employee Name:', 'Alexandra Rivera'),
        ('Department:', 'Product Development'),
        ('Position:', 'Senior Software Engineer'),
        ('Review Period:', 'April 2025 - March 2026'),
        ('Reviewing Manager:', 'David Okonkwo'),
    ]

    for i, (label, value) in enumerate(fields):
        # Label cell - bold
        cell_label = info_table.cell(i, 0)
        cell_label.text = ''
        run_label = cell_label.paragraphs[0].add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(11)

        # Value cell
        cell_value = info_table.cell(i, 1)
        cell_value.text = ''
        run_value = cell_value.paragraphs[0].add_run(value)
        run_value.font.size = Pt(11)

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()  # spacer

    # --- Review Purpose Statement ---
    purpose_heading = doc.add_heading('Review Purpose', level=1)

    purpose_text = (
        'This performance review is intended to provide a comprehensive evaluation of '
        'the employee\'s contributions, competencies, and professional growth over the '
        'review period. The assessment incorporates self-evaluation, managerial observations, '
        'and aggregated peer feedback to ensure a balanced and thorough appraisal. '
        'Results from this review will inform compensation adjustments, development planning, '
        'and career progression discussions.'
    )
    para = doc.add_paragraph(purpose_text)
    para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()  # spacer

    # --- Placeholder note for assessor ---
    note = doc.add_paragraph()
    note.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_note = note.add_run(
        'Note: The competency assessment table, weighted scoring calculations, '
        'and performance level legend will be added to this document as part of '
        'the review process.'
    )
    run_note.italic = True
    run_note.font.size = Pt(10)
    run_note.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
