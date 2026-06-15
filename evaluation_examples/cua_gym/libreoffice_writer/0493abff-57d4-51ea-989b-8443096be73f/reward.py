"""
Reward Script: Performance review document with embedded scoring calculations
Task ID: writer_hr_061
Domain: libreoffice_writer
Scoring:
  C1 (0.25) - Competency assessment table exists with correct structure (12 rows x 5 cols, correct headers)
  C2 (0.20) - 10 competency areas with weight percentages summing to 100%
  C3 (0.15) - Three assessment columns populated with numeric scores in range 1-5
  C4 (0.15) - Weighted total row at bottom of competency table
  C5 (0.10) - Assessment scoring guide section present with 1-5 scale descriptions
  C6 (0.15) - Performance level classification table with 4 levels (Exceeds, Meets, Below, Unsatisfactory)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_061'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 2 tables (initial has only 1)
    tables = doc.tables
    if len(tables) < 2:
        print(f"FAIL: Document has only {len(tables)} table(s). Need at least 2 (employee info + competency).")
        print("REWARD: 0.0")
        return 0.0

    # ---- Component 1: Competency assessment table structure (0.25 pts) ----
    # The competency table should be the 2nd table, with header row + 10 competencies + 1 total = 12 rows
    # Columns: Competency Area, Weight (%), Self-Assessment, Manager Assessment, Peer Average = 5 cols
    try:
        comp_table = tables[1]  # 2nd table (index 1)
        num_rows = len(comp_table.rows)
        num_cols = len(comp_table.columns)

        # Check header row
        header_cells = [comp_table.cell(0, c).text.strip().lower() for c in range(min(num_cols, 5))]
        has_competency_header = any('competency' in h for h in header_cells)
        has_weight_header = any('weight' in h for h in header_cells)
        has_self_header = any('self' in h for h in header_cells)
        has_manager_header = any('manager' in h for h in header_cells)
        has_peer_header = any('peer' in h for h in header_cells)

        headers_ok = all([has_competency_header, has_weight_header, has_self_header, has_manager_header, has_peer_header])
        structure_ok = num_rows >= 12 and num_cols >= 5

        if structure_ok and headers_ok:
            print(f"PASS: Component 1 - Competency table has {num_rows} rows x {num_cols} cols with correct headers (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - Table structure: {num_rows}x{num_cols} (need >=12x5), headers_ok={headers_ok}")
            # Partial: if table exists with roughly right dimensions but headers off
            if num_rows >= 10 and num_cols >= 4:
                total_score += 0.1
                print(f"  Partial credit: 0.1 pts for approximate structure")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # ---- Component 2: 10 competency areas with weights summing to 100% (0.20 pts) ----
    try:
        comp_table = tables[1]
        num_rows = len(comp_table.rows)
        weight_col = None
        # Find weight column index
        for c in range(len(comp_table.columns)):
            h = comp_table.cell(0, c).text.strip().lower()
            if 'weight' in h:
                weight_col = c
                break

        if weight_col is not None:
            weights = []
            competency_names = []
            # Data rows are rows 1 to (num_rows - 2), last row is total
            for r in range(1, num_rows - 1):
                name = comp_table.cell(r, 0).text.strip()
                weight_text = comp_table.cell(r, weight_col).text.strip().replace('%', '')
                if name and weight_text:
                    try:
                        w = float(weight_text)
                        weights.append(w)
                        competency_names.append(name)
                    except ValueError:
                        pass

            weight_sum = sum(weights)
            has_10_competencies = len(competency_names) >= 10
            weights_sum_100 = abs(weight_sum - 100) < 1.0

            if has_10_competencies and weights_sum_100:
                print(f"PASS: Component 2 - {len(competency_names)} competencies with weights summing to {weight_sum}% (0.20 pts)")
                total_score += 0.20
            elif has_10_competencies:
                print(f"FAIL: Component 2 - {len(competency_names)} competencies but weights sum to {weight_sum}% (not 100%)")
                total_score += 0.1
            else:
                print(f"FAIL: Component 2 - Only {len(competency_names)} competencies (need 10), weights sum={weight_sum}%")
        else:
            print(f"FAIL: Component 2 - No weight column found")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # ---- Component 3: Three assessment columns with numeric scores 1-5 (0.15 pts) ----
    try:
        comp_table = tables[1]
        num_cols = len(comp_table.columns)
        num_rows = len(comp_table.rows)

        # Find the three assessment columns by header
        assessment_cols = []
        for c in range(num_cols):
            h = comp_table.cell(0, c).text.strip().lower()
            if 'self' in h or 'manager' in h or 'peer' in h:
                assessment_cols.append(c)

        if len(assessment_cols) >= 3:
            valid_scores = 0
            total_cells = 0
            for r in range(1, num_rows - 1):  # skip header and total row
                for c in assessment_cols:
                    cell_text = comp_table.cell(r, c).text.strip()
                    if cell_text:
                        try:
                            val = float(cell_text)
                            if 1.0 <= val <= 5.0:
                                valid_scores += 1
                            total_cells += 1
                        except ValueError:
                            total_cells += 1

            # Expect at least 10 competencies * 3 columns = 30 scores
            if valid_scores >= 25:
                print(f"PASS: Component 3 - {valid_scores} valid scores in range 1-5 across 3 assessment columns (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 - Only {valid_scores} valid scores (need >= 25)")
        else:
            print(f"FAIL: Component 3 - Found only {len(assessment_cols)} assessment columns (need 3)")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # ---- Component 4: Weighted total row at bottom (0.15 pts) ----
    try:
        comp_table = tables[1]
        num_rows = len(comp_table.rows)
        last_row_text = comp_table.cell(num_rows - 1, 0).text.strip().lower()

        has_total_label = 'total' in last_row_text or 'weighted' in last_row_text

        if has_total_label:
            # Check that the total row has numeric values in assessment columns
            total_values = []
            for c in range(1, len(comp_table.columns)):
                val_text = comp_table.cell(num_rows - 1, c).text.strip().replace('%', '')
                if val_text:
                    try:
                        total_values.append(float(val_text))
                    except ValueError:
                        pass

            if len(total_values) >= 3:
                print(f"PASS: Component 4 - Weighted total row found with values {total_values} (0.15 pts)")
                total_score += 0.15
            elif len(total_values) >= 1:
                print(f"FAIL: Component 4 - Total row has label but only {len(total_values)} numeric values (partial: 0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 4 - Total row label found but no numeric values")
        else:
            print(f"FAIL: Component 4 - Last row label is '{last_row_text}', expected 'total' or 'weighted'")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # ---- Component 5: Assessment scoring guide section (0.10 pts) ----
    # Should have a section explaining the 1-5 scale with descriptions
    try:
        all_text = '\n'.join(p.text for p in doc.paragraphs).lower()

        scoring_heading_count = sum(
            1 for p in doc.paragraphs
            if (('scoring' in p.text.strip().lower() or 'assessment' in p.text.strip().lower() or 'rating' in p.text.strip().lower())
                and ('guide' in p.text.strip().lower() or 'scale' in p.text.strip().lower() or 'criteria' in p.text.strip().lower()))
        )
        has_scoring_heading = scoring_heading_count > 0

        # Check for scale descriptions (1-5 with labels)
        has_scale_1 = ('unsatisfactory' in all_text or '1 -' in all_text or '1 =' in all_text)
        has_scale_5 = ('outstanding' in all_text or 'exceptional' in all_text or '5 -' in all_text or '5 =' in all_text)
        has_meets = 'meets' in all_text
        has_exceeds = 'exceeds' in all_text

        scale_descriptions = sum([has_scale_1, has_scale_5, has_meets, has_exceeds])

        if has_scoring_heading and scale_descriptions >= 3:
            print(f"PASS: Component 5 - Scoring guide section with scale descriptions found (0.10 pts)")
            total_score += 0.10
        elif scale_descriptions >= 3:
            # Guide text exists but heading may be slightly different
            print(f"PASS: Component 5 - Scale descriptions found ({scale_descriptions}/4 keywords) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 - Scoring guide: heading={has_scoring_heading}, scale_keywords={scale_descriptions}/4")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # ---- Component 6: Performance level classification table (0.15 pts) ----
    # Should have a table mapping total scores to performance levels
    try:
        perf_table = None
        for t_idx, t in enumerate(tables):
            if t_idx <= 1:
                continue  # Skip employee info and competency tables
            # Check if this table has performance level content
            first_cell = t.cell(0, 0).text.strip().lower()
            if 'performance' in first_cell or 'level' in first_cell:
                perf_table = t
                break

        if perf_table is None and len(tables) >= 3:
            # Try the 3rd table regardless
            perf_table = tables[2]

        if perf_table is not None:
            perf_rows = len(perf_table.rows)
            all_perf_text = ' '.join(
                perf_table.cell(r, c).text.strip().lower()
                for r in range(perf_rows)
                for c in range(len(perf_table.columns))
            )

            has_exceeds = 'exceeds' in all_perf_text
            has_meets = 'meets' in all_perf_text
            has_below = 'below' in all_perf_text
            has_unsatisfactory = 'unsatisfactory' in all_perf_text

            level_count = sum([has_exceeds, has_meets, has_below, has_unsatisfactory])

            if level_count >= 4:
                print(f"PASS: Component 6 - Performance level table with all 4 levels found (0.15 pts)")
                total_score += 0.15
            elif level_count >= 2:
                partial = 0.075
                print(f"FAIL: Component 6 - Only {level_count}/4 performance levels found (partial: {partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 - Performance level table missing or has only {level_count}/4 levels")
        else:
            print(f"FAIL: Component 6 - No performance level classification table found")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
