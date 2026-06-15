"""
Initial Setup: Document with 3 chapters and 6 images with manually typed figure numbers.
Task ID: writer_tech_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image as PILImage
import io

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_064'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_placeholder_image(path, width=400, height=250, color=(200, 210, 230), label=""):
    """Create a simple placeholder image."""
    img = PILImage.new('RGB', (width, height), color)
    # Draw a simple border
    pixels = img.load()
    for x in range(width):
        pixels[x, 0] = (100, 100, 100)
        pixels[x, height - 1] = (100, 100, 100)
    for y in range(height):
        pixels[0, y] = (100, 100, 100)
        pixels[width - 1, y] = (100, 100, 100)
    img.save(path)
    return path


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Network Infrastructure Modernization Plan', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This document outlines the comprehensive plan for upgrading the corporate '
        'network infrastructure across all regional offices. The project spans three '
        'major phases covering hardware deployment, software configuration, and '
        'security hardening procedures.'
    )
    intro.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        'All diagrams and screenshots referenced in this document reflect the current '
        'state of the network as of Q1 2026. Technical specifications are subject to '
        'change based on vendor availability and budget approvals.'
    )

    # ============================================================
    # Chapter 1: Hardware Infrastructure
    # ============================================================
    doc.add_page_break()
    h1 = doc.add_heading('Chapter 1: Hardware Infrastructure', level=1)

    doc.add_paragraph(
        'The hardware infrastructure upgrade focuses on replacing legacy switches and '
        'routers with next-generation equipment capable of supporting 10 Gbps backbone '
        'connections. The deployment will occur in three waves, starting with the '
        'headquarters data center in Austin, Texas.'
    )

    doc.add_heading('1.1 Core Switch Topology', level=2)
    doc.add_paragraph(
        'The core switching layer will be upgraded from Catalyst 6500 series to Nexus '
        '9300-FX3 switches. Each distribution closet will maintain dual uplinks to the '
        'core layer for redundancy. The following diagram shows the proposed topology '
        'for the Austin headquarters facility.'
    )

    # Image 1 - topology diagram
    img1_path = f'{WORKDIR}/img_topology.png'
    create_placeholder_image(img1_path, 500, 300, (180, 200, 220), "Topology")
    doc.add_picture(img1_path, width=Inches(5.0))
    caption1 = doc.add_paragraph('Figure 1-1: Core switch topology for Austin headquarters')
    caption1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in caption1.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph(
        'The topology employs a leaf-spine architecture with four spine switches and '
        'twelve leaf switches distributed across floors 2 through 7. Each leaf switch '
        'provides 48 access ports with PoE+ capability for wireless access points and '
        'VoIP handsets.'
    )

    doc.add_heading('1.2 Server Rack Layout', level=2)
    doc.add_paragraph(
        'The server room redesign consolidates 14 half-populated racks into 8 fully '
        'utilized racks with improved airflow management. Hot aisle containment will be '
        'installed between rows B and C to achieve target PUE of 1.35.'
    )

    # Image 2 - rack layout
    img2_path = f'{WORKDIR}/img_rack_layout.png'
    create_placeholder_image(img2_path, 450, 350, (190, 210, 190), "Rack Layout")
    doc.add_picture(img2_path, width=Inches(4.5))
    caption2 = doc.add_paragraph('Figure 1-2: Server rack layout with hot aisle containment')
    caption2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in caption2.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph(
        'Power distribution units rated at 30A/208V will be installed in each rack, '
        'with A+B feed redundancy from separate UPS systems. The estimated total power '
        'draw per rack is 8.5 kW under peak load conditions.'
    )

    # ============================================================
    # Chapter 2: Software Configuration
    # ============================================================
    doc.add_page_break()
    h2 = doc.add_heading('Chapter 2: Software Configuration', level=1)

    doc.add_paragraph(
        'Software configuration encompasses operating system deployment, network '
        'management tools, and monitoring infrastructure. All systems will be managed '
        'through a centralized Ansible control node with Git-based configuration '
        'version control.'
    )

    doc.add_heading('2.1 VLAN Architecture', level=2)
    doc.add_paragraph(
        'The VLAN scheme has been redesigned to support micro-segmentation requirements '
        'from the security team. Each department receives a dedicated /24 subnet within '
        'the 10.128.0.0/16 address space. Inter-VLAN routing is handled by the core '
        'switches using hardware-accelerated Layer 3 forwarding.'
    )

    # Image 3 - VLAN diagram
    img3_path = f'{WORKDIR}/img_vlan.png'
    create_placeholder_image(img3_path, 480, 280, (220, 200, 180), "VLAN")
    doc.add_picture(img3_path, width=Inches(4.8))
    caption3 = doc.add_paragraph('Figure 2-1: VLAN segmentation diagram with subnet allocations')
    caption3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in caption3.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph(
        'Access control lists enforce traffic isolation between VLANs. The management '
        'VLAN (VLAN 999) is restricted to authorized administrators via 802.1X '
        'authentication with RADIUS backend. Guest wireless traffic is isolated on '
        'VLAN 500 with captive portal integration.'
    )

    doc.add_heading('2.2 Monitoring Dashboard', level=2)
    doc.add_paragraph(
        'A Grafana-based monitoring stack will provide real-time visibility into network '
        'performance metrics. Prometheus exporters deployed on each managed switch will '
        'collect interface statistics, CPU utilization, and memory consumption at '
        '30-second intervals.'
    )

    # Image 4 - monitoring dashboard
    img4_path = f'{WORKDIR}/img_dashboard.png'
    create_placeholder_image(img4_path, 520, 300, (200, 190, 210), "Dashboard")
    doc.add_picture(img4_path, width=Inches(5.2))
    caption4 = doc.add_paragraph('Figure 2-2: Network monitoring dashboard overview')
    caption4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in caption4.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph(
        'Alert thresholds are configured for interface utilization above 80%, packet '
        'loss exceeding 0.1%, and latency spikes beyond 5ms on the backbone links. '
        'Notifications route through PagerDuty to the on-call network engineering team.'
    )

    # ============================================================
    # Chapter 3: Security Hardening
    # ============================================================
    doc.add_page_break()
    h3 = doc.add_heading('Chapter 3: Security Hardening', level=1)

    doc.add_paragraph(
        'Security hardening addresses both perimeter defenses and internal segmentation. '
        'The zero-trust model requires all devices to authenticate before accessing '
        'any network resources, regardless of physical location within the campus.'
    )

    doc.add_heading('3.1 Firewall Rule Matrix', level=2)
    doc.add_paragraph(
        'The firewall policy has been redesigned using a deny-all default stance with '
        'explicit permit rules for authorized traffic flows. The rule set was developed '
        'in collaboration with the application teams to document every required '
        'communication path.'
    )

    # Image 5 - firewall matrix
    img5_path = f'{WORKDIR}/img_firewall.png'
    create_placeholder_image(img5_path, 500, 320, (210, 195, 195), "Firewall")
    doc.add_picture(img5_path, width=Inches(5.0))
    caption5 = doc.add_paragraph('Figure 3-1: Firewall rule matrix between security zones')
    caption5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in caption5.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph(
        'The DMZ zone hosts all externally accessible services including the corporate '
        'website, VPN concentrators, and email gateways. Outbound traffic from the '
        'internal zone passes through a web proxy with SSL inspection for content '
        'filtering and data loss prevention.'
    )

    doc.add_heading('3.2 Intrusion Detection Layout', level=2)
    doc.add_paragraph(
        'Network-based intrusion detection sensors are deployed at each trust boundary. '
        'Traffic mirroring from the core switches feeds dedicated Suricata instances '
        'running custom rule sets tuned for the organization\'s threat profile.'
    )

    # Image 6 - IDS layout
    img6_path = f'{WORKDIR}/img_ids.png'
    create_placeholder_image(img6_path, 460, 290, (195, 215, 200), "IDS")
    doc.add_picture(img6_path, width=Inches(4.6))
    caption6 = doc.add_paragraph('Figure 3-2: Intrusion detection sensor placement diagram')
    caption6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in caption6.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph(
        'The central SIEM aggregates alerts from all IDS sensors, correlating events '
        'across network zones to identify lateral movement patterns. Automated playbooks '
        'trigger port isolation for hosts exhibiting indicators of compromise, with a '
        'mean time to containment target of under 15 minutes.'
    )

    # --- Conclusion ---
    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The network modernization project represents a significant investment in the '
        'organization\'s digital infrastructure. Completion of all three phases is '
        'projected for Q3 2026, with measurable improvements in throughput, reliability, '
        'and security posture expected within the first quarter of full operation.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up placeholder images
    for img_file in ['img_topology.png', 'img_rack_layout.png', 'img_vlan.png',
                     'img_dashboard.png', 'img_firewall.png', 'img_ids.png']:
        try:
            os.remove(f'{WORKDIR}/{img_file}')
        except OSError:
            pass

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
