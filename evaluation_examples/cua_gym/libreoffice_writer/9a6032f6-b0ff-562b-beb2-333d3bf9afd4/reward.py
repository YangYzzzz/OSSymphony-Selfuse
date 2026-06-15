"""
Reward Script: Professional resume layout using tables with hidden borders
Task ID: writer_rd_035
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20) — Table exists with 2 columns and multiple rows
  Component 2 (0.20) — All table borders are hidden (set to 'none')
  Component 3 (0.25) — Left column has bold section headers (Experience, Education, Skills)
  Component 4 (0.15) — Right column has corresponding detail content
  Component 5 (0.20) — Separator rows with colored shading exist
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_035'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document — resume layout requires a table")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table exists with 2 columns and multiple rows (0.20 points)
    try:
        if num_cols == 2 and num_rows >= 3:
            print(f"PASS: Component 1 — Table has {num_rows} rows and {num_cols} columns (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected 2 cols and >=3 rows, found {num_cols} cols and {num_rows} rows")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All table borders are hidden/none (0.20 points)
    try:
        tblPr = table._tbl.find(qn('w:tblPr'))
        borders_hidden = False
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                border_types = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
                all_none = True
                for bt in border_types:
                    border_el = tblBorders.find(qn(f'w:{bt}'))
                    if border_el is not None:
                        val = border_el.get(qn('w:val'))
                        sz = border_el.get(qn('w:sz'))
                        # 'none' or 'nil' or size=0 means hidden
                        if val not in ('none', 'nil') and sz != '0':
                            all_none = False
                            print(f"  Border '{bt}': val={val}, sz={sz} — NOT hidden")
                    # If border element is missing, it defaults to no border (ok)
                borders_hidden = all_none

        if borders_hidden:
            print(f"PASS: Component 2 — All table borders are hidden (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — Table borders are not all hidden")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Left column has bold section headers (Experience, Education, Skills) (0.25 points)
    try:
        required_headers = {'experience', 'education', 'skills'}
        found_headers = set()

        for ri in range(num_rows):
            cell = table.cell(ri, 0)
            cell_text = cell.text.strip().lower()
            for header in required_headers:
                if header in cell_text:
                    # Check if the text is bold
                    has_bold = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip() and run.bold:
                                has_bold = True
                                break
                        if has_bold:
                            break
                    if has_bold:
                        found_headers.add(header)

        if found_headers == required_headers:
            print(f"PASS: Component 3 — All 3 section headers found in bold: {found_headers} (0.25 pts)")
            total_score += 0.25
        elif len(found_headers) >= 2:
            partial = round(0.25 * len(found_headers) / 3, 2)
            print(f"PARTIAL: Component 3 — Found {len(found_headers)}/3 bold headers: {found_headers} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Found only {len(found_headers)}/3 bold section headers in left column: {found_headers}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Right column has corresponding detail content (0.15 points)
    try:
        # Check that right column cells (non-empty rows) have substantial text content
        right_col_content_count = 0
        for ri in range(num_rows):
            cell = table.cell(ri, 1)
            cell_text = cell.text.strip()
            # Count cells with at least 20 chars of content (not just labels)
            if len(cell_text) >= 20:
                right_col_content_count += 1

        if right_col_content_count >= 3:
            print(f"PASS: Component 4 — Right column has {right_col_content_count} content-rich cells (0.15 pts)")
            total_score += 0.15
        elif right_col_content_count >= 1:
            partial = round(0.15 * right_col_content_count / 3, 2)
            print(f"PARTIAL: Component 4 — Right column has {right_col_content_count}/3 content-rich cells ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Right column has no substantial content")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Separator rows with colored shading exist (0.20 points)
    try:
        shaded_rows = 0
        for ri in range(num_rows):
            row_shaded = False
            for ci in range(num_cols):
                tc = table.cell(ri, ci)._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill.lower() not in ('auto', 'ffffff', ''):
                            row_shaded = True
                            break
            # Only count as separator if the row is mostly empty (separator purpose)
            row_text = ''.join(table.cell(ri, ci).text.strip() for ci in range(num_cols))
            if row_shaded and len(row_text) < 5:
                shaded_rows += 1

        if shaded_rows >= 2:
            print(f"PASS: Component 5 — Found {shaded_rows} shaded separator rows (0.20 pts)")
            total_score += 0.20
        elif shaded_rows >= 1:
            print(f"PARTIAL: Component 5 — Found {shaded_rows} shaded separator row (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — No shaded separator rows found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
