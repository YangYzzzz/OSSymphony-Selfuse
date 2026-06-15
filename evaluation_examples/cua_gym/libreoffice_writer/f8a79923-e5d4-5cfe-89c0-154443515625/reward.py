"""
Reward Script: Mail merge thank-you letter for professors
Task ID: writer_mt_044
Domain: libreoffice_writer
Scoring:
  Component 1: Letter greeting with "Dear Professor" (0.15 pts)
  Component 2: MERGEFIELD ProfessorName present (0.25 pts)
  Component 3: MERGEFIELD Department present (0.20 pts)
  Component 4: MERGEFIELD University present (0.20 pts)
  Component 5: Closing signature with "Sincerely" (0.20 pts)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_044'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Checks that the document contains a thank-you letter with proper
    mail merge fields (ProfessorName, Department, University) and a
    closing signature.
    """
    total_score = 0.0

    # Load the document
    try:
        from docx import Document
        from lxml import etree
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have some content
    all_text = ' '.join(p.text for p in doc.paragraphs).strip()
    if not all_text:
        print("FAIL: Document is empty — no content found")
        print("REWARD: 0.0")
        return 0.0

    # Helper: extract all MERGEFIELD names from document XML
    def get_merge_field_names(document):
        """Parse the document XML to find all MERGEFIELD instrText entries."""
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        field_names = set()
        for para in document.paragraphs:
            instr_elements = para._element.findall('.//w:instrText', ns)
            for instr in instr_elements:
                text = instr.text or ''
                text = text.strip()
                if text.upper().startswith('MERGEFIELD'):
                    # Extract field name: "MERGEFIELD ProfessorName" -> "ProfessorName"
                    parts = text.split()
                    if len(parts) >= 2:
                        field_names.add(parts[1])
        return field_names

    merge_fields = get_merge_field_names(doc)
    print(f"INFO: Found merge fields: {merge_fields}")
    print(f"INFO: Full text preview: {all_text[:200]}")

    # Component 1: Letter greeting with "Dear Professor" (0.15 points)
    # This checks that the document is structured as a letter with the
    # proper salutation. The initial doc is blank, so this only passes on golden.
    try:
        first_para_text = doc.paragraphs[0].text if doc.paragraphs else ''
        if 'dear professor' in first_para_text.lower():
            print(f"PASS: Component 1 — Greeting found: '{first_para_text}' (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Expected greeting with 'Dear Professor', found: '{first_para_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: MERGEFIELD ProfessorName present (0.25 points)
    # The task requires ProfessorName as a proper mail merge field, not just
    # plain text placeholder. Must be a real MERGEFIELD in the document XML.
    try:
        if 'ProfessorName' in merge_fields:
            print(f"PASS: Component 2 — MERGEFIELD 'ProfessorName' found (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — MERGEFIELD 'ProfessorName' not found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: MERGEFIELD Department present (0.20 points)
    try:
        if 'Department' in merge_fields:
            print(f"PASS: Component 3 — MERGEFIELD 'Department' found (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — MERGEFIELD 'Department' not found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: MERGEFIELD University present (0.20 points)
    try:
        if 'University' in merge_fields:
            print(f"PASS: Component 4 — MERGEFIELD 'University' found (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — MERGEFIELD 'University' not found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Closing signature with "Sincerely" (0.20 points)
    # The task says the letter should be signed "Sincerely, [Student Name]".
    # Check that at least one paragraph contains "sincerely" (case insensitive).
    try:
        has_sincerely = any('sincerely' in p.text.lower() for p in doc.paragraphs)
        if has_sincerely:
            print(f"PASS: Component 5 — Closing 'Sincerely' found (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 5 — No paragraph contains 'Sincerely' closing")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
