"""
Initial Setup: Create a Writer document with equal margins and realistic content.
Task ID: writer_tech_074
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set equal margins on all sides (2.54 cm = 1 inch, standard default)
    for section in doc.sections:
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    # --- Title ---
    title = doc.add_heading('Network Infrastructure Migration Plan', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle / metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Prepared by: Infrastructure Engineering Team')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = meta2.add_run('Version 2.3 — March 2025')
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines the comprehensive plan for migrating our legacy '
        'network infrastructure from the current on-premises data center at the '
        'Portland facility to the new hybrid cloud architecture. The migration '
        'will affect approximately 340 servers, 12 network switches, and over '
        '50 virtual LANs currently serving 2,800 employees across six regional offices.'
    )
    doc.add_paragraph(
        'The project timeline spans 18 months, with an estimated budget of $4.2 million. '
        'Phase 1 (core routing) is scheduled to begin on April 15, 2025, with final '
        'cutover targeted for September 30, 2026. Risk mitigation strategies include '
        'parallel operation during transition windows and automated failover testing.'
    )

    # --- Section 2: Current Architecture ---
    doc.add_heading('2. Current Architecture Assessment', level=1)
    doc.add_paragraph(
        'The existing infrastructure relies on a three-tier network topology deployed '
        'in 2018. Core routing is handled by four Cisco Nexus 9500 switches operating '
        'in a redundant pair configuration. The distribution layer comprises 12 Catalyst '
        '9300 series switches providing inter-VLAN routing and QoS policy enforcement.'
    )

    # Sub-section
    doc.add_heading('2.1 Hardware Inventory', level=2)

    # Table: Hardware inventory
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['Component', 'Model', 'Quantity', 'End of Support']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    inventory = [
        ['Core Switches', 'Cisco Nexus 9500', '4', 'Dec 2027'],
        ['Distribution Switches', 'Catalyst 9300', '12', 'Jun 2026'],
        ['Access Points', 'Meraki MR56', '85', 'Mar 2028'],
        ['Firewalls', 'Palo Alto PA-5260', '2', 'Sep 2026'],
        ['Load Balancers', 'F5 BIG-IP i5800', '2', 'Jan 2027'],
        ['WAN Optimizers', 'Riverbed SteelHead', '6', 'Nov 2025'],
    ]
    for r, row_data in enumerate(inventory, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    doc.add_heading('2.2 Traffic Analysis', level=2)
    doc.add_paragraph(
        'Average daily throughput across the core backbone is 48 Gbps with peak loads '
        'reaching 72 Gbps during the 9:00-11:00 AM window. East-west traffic between '
        'application tiers accounts for 65% of total bandwidth, while north-south '
        'traffic (internet-bound) represents the remaining 35%. VoIP traffic averages '
        '3.2 Gbps with a 99.97% uptime SLA requirement.'
    )

    # --- Section 3: Target Architecture ---
    doc.add_heading('3. Target Architecture', level=1)
    doc.add_paragraph(
        'The proposed hybrid architecture combines a collocated private cloud footprint '
        'with AWS transit gateway peering. Software-defined networking (SD-WAN) will '
        'replace traditional MPLS circuits for branch office connectivity, reducing '
        'monthly WAN costs by an estimated 40%.'
    )

    doc.add_heading('3.1 Cloud Components', level=2)
    items = [
        'AWS Direct Connect (10 Gbps dedicated) for primary cloud path',
        'Azure ExpressRoute (5 Gbps) for disaster recovery failover',
        'Cloudflare Magic Transit for DDoS mitigation and edge caching',
        'HashiCorp Consul for service mesh and DNS-based load balancing',
        'Terraform-managed infrastructure as code for all cloud resources',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # --- Section 4: Migration Phases ---
    doc.add_heading('4. Migration Phases', level=1)

    doc.add_heading('4.1 Phase 1 — Core Routing (Weeks 1-12)', level=2)
    doc.add_paragraph(
        'Deploy spine-leaf fabric in the new data center wing using Arista 7280R3 '
        'switches. Establish BGP peering with existing core and validate ECMP load '
        'balancing. Target completion: July 2025.'
    )

    doc.add_heading('4.2 Phase 2 — Application Migration (Weeks 13-36)', level=2)
    doc.add_paragraph(
        'Migrate application workloads in priority order: (1) stateless web frontends, '
        '(2) containerized microservices, (3) database clusters with replication, '
        '(4) legacy monolithic applications requiring re-platforming. Each migration '
        'window will include 72-hour parallel operation and automated smoke testing.'
    )

    doc.add_heading('4.3 Phase 3 — Decommission (Weeks 37-52)', level=2)
    doc.add_paragraph(
        'Systematically power down legacy equipment following the reverse dependency '
        'order. All decommissioned hardware will be certified for data destruction per '
        'NIST SP 800-88 guidelines before disposal. Estimated recovery value from '
        'hardware resale: $180,000.'
    )

    # --- Section 5: Risk Assessment ---
    doc.add_heading('5. Risk Assessment', level=1)

    risk_table = doc.add_table(rows=5, cols=4)
    risk_table.style = 'Table Grid'
    risk_headers = ['Risk Category', 'Likelihood', 'Impact', 'Mitigation']
    for i, h in enumerate(risk_headers):
        cell = risk_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    risks = [
        ['Service disruption during cutover', 'Medium', 'High',
         'Parallel operation with automated failback'],
        ['Vendor delivery delays', 'Low', 'Medium',
         'Pre-ordered equipment with 90-day buffer'],
        ['Staff skill gaps in SD-WAN', 'High', 'Medium',
         'Training program starting Q2 2025'],
        ['Budget overrun', 'Medium', 'Medium',
         '15% contingency reserve allocated'],
    ]
    for r, row_data in enumerate(risks, 1):
        for c, val in enumerate(row_data):
            risk_table.cell(r, c).text = val

    doc.add_paragraph()

    # --- Section 6: Budget ---
    doc.add_heading('6. Budget Summary', level=1)
    doc.add_paragraph(
        'Total approved budget: $4,200,000. Hardware procurement accounts for '
        '$2,100,000 (50%), professional services $840,000 (20%), cloud infrastructure '
        '$630,000 (15%), training and certification $315,000 (7.5%), and contingency '
        'reserve $315,000 (7.5%).'
    )

    # --- Section 7: Approval ---
    doc.add_heading('7. Approvals', level=1)
    doc.add_paragraph(
        'This plan requires sign-off from the CTO, VP of Infrastructure, '
        'Information Security Officer, and the Finance Committee before Phase 1 '
        'procurement can proceed.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
