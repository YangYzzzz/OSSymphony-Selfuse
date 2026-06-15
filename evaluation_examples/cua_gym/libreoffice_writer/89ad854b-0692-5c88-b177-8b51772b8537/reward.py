"""
Reward Script: Fax Cover Sheet in LibreOffice Writer
Task ID: writer_wf_080
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): FAX title - centered, bold, 30pt
  Component 2 (0.25): Info table - 7 rows with correct field labels
  Component 3 (0.20): Checkbox line - 4 checkboxes (Urgent, For Review, Please Comment, Please Reply)
  Component 4 (0.15): Comments section - bold "Comments:" label with blank lines
  Component 5 (0.15): Confidentiality notice - italic, small font text at bottom
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_080'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify fax cover sheet creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = ' '.join(p.text for p in doc.paragraphs)

    # Component 1: FAX title - centered, bold, 30pt (0.25 points)
    try:
        fax_found = False
        for para in doc.paragraphs:
            if para.text.strip().upper() == 'FAX':
                # Check centering
                is_centered = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                # Check bold and size in runs
                has_bold = False
                has_30pt = False
                for run in para.runs:
                    if run.text.strip():
                        if run.font.bold:
                            has_bold = True
                        if run.font.size and abs(run.font.size.pt - 30.0) < 2.0:
                            has_30pt = True
                if is_centered and has_bold and has_30pt:
                    print(f"PASS: Component 1 — FAX title is centered, bold, 30pt (0.25 pts)")
                    total_score += 0.25
                    fax_found = True
                elif is_centered or has_bold or has_30pt:
                    # Partial: at least some formatting correct
                    partial = 0.0
                    if is_centered:
                        partial += 0.08
                    if has_bold:
                        partial += 0.08
                    if has_30pt:
                        partial += 0.09
                    print(f"PARTIAL: Component 1 — FAX title found, centered={is_centered}, bold={has_bold}, 30pt={has_30pt} ({partial:.2f} pts)")
                    total_score += partial
                    fax_found = True
                else:
                    print(f"FAIL: Component 1 — FAX text found but missing formatting: centered={is_centered}, bold={has_bold}, 30pt={has_30pt}")
                    fax_found = True
                break
        if not fax_found:
            print("FAIL: Component 1 — No 'FAX' title paragraph found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Info table with 7 fields (0.25 points)
    try:
        expected_fields = ['to', 'from', 'fax number', 'phone number', 'date', 'number of pages', 're']
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            found_fields = []
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip().lower()
                    for ef in expected_fields:
                        if ef in cell_text and ef not in found_fields:
                            found_fields.append(ef)
            match_count = len(found_fields)
            if match_count >= 7:
                print(f"PASS: Component 2 — Table has all 7 required fields ({match_count}/7) (0.25 pts)")
                total_score += 0.25
            elif match_count >= 4:
                pts = round(0.25 * match_count / 7, 2)
                print(f"PARTIAL: Component 2 — Table has {match_count}/7 fields: {found_fields} ({pts} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 2 — Table has only {match_count}/7 fields: {found_fields}")
        else:
            print("FAIL: Component 2 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Checkbox line with 4 options (0.20 points)
    try:
        checkbox_labels = ['urgent', 'for review', 'please comment', 'please reply']
        checkbox_char = '\u2610'  # ☐
        found_checkboxes = 0
        for para in doc.paragraphs:
            para_lower = para.text.lower()
            if checkbox_char in para.text or any(label in para_lower for label in checkbox_labels):
                for label in checkbox_labels:
                    if label in para_lower:
                        found_checkboxes += 1
                break  # Only need one paragraph with checkboxes

        if found_checkboxes >= 4:
            print(f"PASS: Component 3 — All 4 checkbox options found (0.20 pts)")
            total_score += 0.20
        elif found_checkboxes >= 2:
            pts = round(0.20 * found_checkboxes / 4, 2)
            print(f"PARTIAL: Component 3 — {found_checkboxes}/4 checkbox options found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 3 — Only {found_checkboxes}/4 checkbox options found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Comments section (0.15 points)
    try:
        comments_found = False
        comments_bold = False
        has_blank_lines = False
        for i, para in enumerate(doc.paragraphs):
            if 'comments' in para.text.lower().strip().rstrip(':'):
                comments_found = True
                # Check if bold
                for run in para.runs:
                    if run.text.strip() and run.font.bold:
                        comments_bold = True
                # Check for blank/underline lines after
                remaining = doc.paragraphs[i+1:i+8] if i+1 < len(doc.paragraphs) else []
                blank_or_underline_count = 0
                for p in remaining:
                    t = p.text.strip()
                    if not t or all(c == '_' for c in t):
                        blank_or_underline_count += 1
                if blank_or_underline_count >= 2:
                    has_blank_lines = True
                break

        if comments_found and comments_bold and has_blank_lines:
            print(f"PASS: Component 4 — Comments section with bold label and blank lines (0.15 pts)")
            total_score += 0.15
        elif comments_found and (comments_bold or has_blank_lines):
            print(f"PARTIAL: Component 4 — Comments found, bold={comments_bold}, blank_lines={has_blank_lines} (0.08 pts)")
            total_score += 0.08
        elif comments_found:
            print(f"PARTIAL: Component 4 — Comments label found but no bold/blank lines (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — No 'Comments' section found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Confidentiality notice - italic, small font (0.15 points)
    try:
        confidential_found = False
        is_italic = False
        is_small = False
        for para in doc.paragraphs:
            if 'confidentiality' in para.text.lower() or 'confidential' in para.text.lower():
                confidential_found = True
                for run in para.runs:
                    if run.text.strip():
                        if run.font.italic:
                            is_italic = True
                        if run.font.size and run.font.size.pt < 11.0:
                            is_small = True
                break

        if confidential_found and is_italic and is_small:
            print(f"PASS: Component 5 — Confidentiality notice in italic, small font (0.15 pts)")
            total_score += 0.15
        elif confidential_found and (is_italic or is_small):
            print(f"PARTIAL: Component 5 — Confidentiality found, italic={is_italic}, small={is_small} (0.08 pts)")
            total_score += 0.08
        elif confidential_found:
            print(f"PARTIAL: Component 5 — Confidentiality text found but no italic/small formatting (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No confidentiality notice found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
