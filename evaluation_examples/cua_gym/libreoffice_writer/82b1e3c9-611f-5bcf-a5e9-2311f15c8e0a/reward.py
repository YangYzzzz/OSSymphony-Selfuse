"""
Reward Script: CD/DVD Label Template Creation
Task ID: writer_lec_058
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Document contains exactly 2 tables (2 labels per sheet)
  Component 2 (0.30): Both tables contain 'Annual Report 2025' as centered text
  Component 3 (0.20): Label text is centered within table cells
  Component 4 (0.25): Document contains CD/DVD label template info (Avery 8692 / 117mm)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_058'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify CD/DVD label template creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_tables = len(doc.tables)
    num_paras = len(doc.paragraphs)

    # Precondition: document must have content (not blank)
    if num_tables == 0 and num_paras == 0:
        print("FAIL: Document is blank — no tables or paragraphs found")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Document has 2 tables representing 2 labels per sheet (0.25 points)
    try:
        if num_tables >= 2:
            print(f"PASS: Component 1 — Found {num_tables} tables (>= 2 labels per sheet) (0.25 pts)")
            total_score += 0.25
        elif num_tables == 1:
            # Partial: at least one label exists
            print(f"PARTIAL: Component 1 — Found 1 table (expected 2 for 2 labels/sheet) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — No tables found (expected 2 for CD labels)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Both tables contain 'Annual Report 2025' text (0.30 points)
    try:
        labels_with_text = 0
        target_text = "annual report 2025"
        for ti in range(min(num_tables, 2)):
            table = doc.tables[ti]
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text.strip().lower() + " "
            if target_text in table_text.strip():
                labels_with_text += 1

        if labels_with_text >= 2:
            print(f"PASS: Component 2 — Both labels contain 'Annual Report 2025' (0.30 pts)")
            total_score += 0.30
        elif labels_with_text == 1:
            print(f"PARTIAL: Component 2 — Only 1 of 2 labels contains 'Annual Report 2025' (0.15 pts)")
            total_score += 0.15
        else:
            # Also check paragraphs in case labels are done differently (not as tables)
            para_has_text = any(target_text in p.text.lower() for p in doc.paragraphs)
            if para_has_text:
                print(f"PARTIAL: Component 2 — 'Annual Report 2025' found in paragraphs but not in table labels (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 2 — 'Annual Report 2025' not found in any label")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text in table cells is centered (0.20 points)
    try:
        centered_labels = 0
        for ti in range(min(num_tables, 2)):
            table = doc.tables[ti]
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            align = para.paragraph_format.alignment
                            if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                                centered_labels += 1
                                break  # One centered para in cell is enough
                    if centered_labels > ti:
                        break  # Found centered text in this table
                if centered_labels > ti:
                    break

        if centered_labels >= 2:
            print(f"PASS: Component 3 — Text is centered in both labels (0.20 pts)")
            total_score += 0.20
        elif centered_labels == 1:
            print(f"PARTIAL: Component 3 — Text centered in 1 of 2 labels (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — Label text is not centered")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Document contains CD/DVD label template metadata (Avery 8692, 117mm) (0.25 points)
    try:
        all_text = " ".join(p.text.lower() for p in doc.paragraphs)
        # Also include table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text.lower()

        has_cd_dvd = bool(re.search(r'cd[/ ]?dvd|cd label|dvd label|label template', all_text))
        has_avery = bool(re.search(r'avery\s*8692', all_text))
        has_117mm = bool(re.search(r'117\s*mm', all_text))
        has_2_per_sheet = bool(re.search(r'2\s*(labels?\s*(per|/)\s*sheet|per\s*sheet)', all_text))

        indicators = sum([has_cd_dvd, has_avery, has_117mm, has_2_per_sheet])

        if indicators >= 3:
            print(f"PASS: Component 4 — Template info present: cd/dvd={has_cd_dvd}, avery={has_avery}, 117mm={has_117mm}, 2/sheet={has_2_per_sheet} (0.25 pts)")
            total_score += 0.25
        elif indicators >= 2:
            print(f"PARTIAL: Component 4 — Some template info: cd/dvd={has_cd_dvd}, avery={has_avery}, 117mm={has_117mm}, 2/sheet={has_2_per_sheet} (0.15 pts)")
            total_score += 0.15
        elif indicators >= 1:
            print(f"PARTIAL: Component 4 — Minimal template info: cd/dvd={has_cd_dvd}, avery={has_avery}, 117mm={has_117mm}, 2/sheet={has_2_per_sheet} (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — No CD/DVD label template metadata found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
