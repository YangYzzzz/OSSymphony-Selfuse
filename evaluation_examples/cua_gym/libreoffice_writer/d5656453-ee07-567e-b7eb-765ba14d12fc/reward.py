"""
Reward Script: Reorganize HR folder and create index.docx
Task ID: writer_hr_060
Domain: libreoffice_writer
Scoring:
  Component 1: HR_Policies_2026 folder exists AND all 4 policy files moved into it (0.3 pts)
  Component 2: index.docx exists inside HR_Policies_2026 (0.2 pts)
  Component 3: index.docx has Heading 1 containing 'HR Policy Documents' title (0.2 pts)
  Component 4: index.docx has a 2-column table with header row and all 4 policy files listed (0.3 pts)
Total: 1.0
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_060'

DESKTOP = os.path.join(WORKDIR, 'Desktop')
HR_FOLDER = os.path.join(DESKTOP, 'HR_Policies_2026')
INDEX_DOC = os.path.join(HR_FOLDER, 'index.docx')

POLICY_FILES = [
    'attendance_policy.docx',
    'leave_policy.docx',
    'remote_work_policy.docx',
    'travel_expense_policy.docx',
]


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Component 1: HR_Policies_2026 folder exists AND all 4 policy files have been
    # moved into it (they should NOT be at ~/Desktop/ anymore) (0.3 points)
    try:
        folder_exists = os.path.isdir(HR_FOLDER)
        if not folder_exists:
            print(f"FAIL: Component 1 — HR_Policies_2026 folder does not exist at {HR_FOLDER}")
        else:
            # Check all 4 files are inside the folder
            files_in_folder = [
                f for f in POLICY_FILES
                if os.path.isfile(os.path.join(HR_FOLDER, f))
            ]
            # Check original files are NOT at Desktop root
            files_still_at_desktop = [
                f for f in POLICY_FILES
                if os.path.isfile(os.path.join(DESKTOP, f))
            ]
            if len(files_in_folder) == 4 and len(files_still_at_desktop) == 0:
                print(f"PASS: Component 1 — all 4 policy files moved to HR_Policies_2026/ (0.3 pts)")
                total_score += 0.3
            else:
                in_count = len(files_in_folder)
                still_count = len(files_still_at_desktop)
                print(f"FAIL: Component 1 — {in_count}/4 files found in HR_Policies_2026/, "
                      f"{still_count} still at Desktop root. "
                      f"In folder: {files_in_folder}. Still at desktop: {files_still_at_desktop}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: index.docx exists inside HR_Policies_2026 (0.2 points)
    try:
        if os.path.isfile(INDEX_DOC):
            print(f"PASS: Component 2 — index.docx exists at {INDEX_DOC} (0.2 pts)")
            total_score += 0.2
        else:
            # Check folder for any index file variant
            if os.path.isdir(HR_FOLDER):
                contents = os.listdir(HR_FOLDER)
                print(f"FAIL: Component 2 — index.docx not found. Folder contains: {contents}")
            else:
                print(f"FAIL: Component 2 — HR_Policies_2026 folder does not exist")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Components 3 and 4 require loading index.docx
    if not os.path.isfile(INDEX_DOC):
        print(f"SKIP: Components 3 & 4 — index.docx not found, cannot verify content")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    try:
        from docx import Document
        doc = Document(INDEX_DOC)
    except Exception as e:
        print(f"CRITICAL: Cannot load index.docx: {e}")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 3: index.docx has a Heading 1 paragraph containing 'HR Policy Documents' (0.2 points)
    try:
        matching_heading1 = [
            p for p in doc.paragraphs
            if ('Heading 1' in p.style.name) and ('HR Policy Documents' in p.text)
        ]
        if len(matching_heading1) > 0:
            print(f"PASS: Component 3 — Heading 1 found: {matching_heading1[0].text!r} (0.2 pts)")
            total_score += 0.2
        else:
            all_headings = [(p.style.name, p.text) for p in doc.paragraphs if 'Heading' in p.style.name]
            print(f"FAIL: Component 3 — No Heading 1 containing 'HR Policy Documents' found. "
                  f"All headings: {all_headings}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: index.docx has a 2-column table with header row 'File Name' | 'Description'
    # and all 4 policy file names listed (0.3 points)
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 4 — No tables found in index.docx")
        else:
            table = doc.tables[0]
            rows = table.rows
            cols = table.columns

            # Check header row
            if len(rows) < 1 or len(cols) < 2:
                print(f"FAIL: Component 4 — Table has {len(rows)} rows, {len(cols)} cols; expected at least 1 row and 2 cols")
            else:
                header_cell_0 = rows[0].cells[0].text.strip()
                header_cell_1 = rows[0].cells[1].text.strip()
                header_ok = (
                    'File Name' in header_cell_0 or header_cell_0.lower() == 'file name'
                )

                # Gather all file names listed in the table (rows 1+)
                listed_files = []
                for row in rows[1:]:
                    cell_text = row.cells[0].text.strip()
                    if cell_text:
                        listed_files.append(cell_text)

                # Check all 4 policy files are listed (exact name match)
                missing_files = [f for f in POLICY_FILES if f not in listed_files]
                has_all_files = len(missing_files) == 0

                if header_ok and has_all_files:
                    print(f"PASS: Component 4 — Table header OK ({header_cell_0!r} | {header_cell_1!r}), "
                          f"all 4 policy files listed (0.3 pts)")
                    total_score += 0.3
                elif not header_ok:
                    print(f"FAIL: Component 4 — Header row missing 'File Name'. Found: "
                          f"{header_cell_0!r} | {header_cell_1!r}")
                else:
                    print(f"FAIL: Component 4 — Header OK but missing policy files: {missing_files}. "
                          f"Listed: {listed_files}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


verify_task()
