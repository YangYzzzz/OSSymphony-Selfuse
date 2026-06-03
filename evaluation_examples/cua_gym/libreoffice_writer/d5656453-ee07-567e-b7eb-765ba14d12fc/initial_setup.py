"""
Initial Setup: HR folder reorganization task
Task ID: writer_hr_060
Domain: libreoffice_writer

Creates 4 HR policy .docx files on the Desktop with realistic content.
The folder ~/Desktop/HR_Policies_2026/ does NOT exist yet (the agent must create it).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_hr_060'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_attendance_policy():
    path = f'{WORKDIR}/attendance_policy.docx'
    doc = Document()
    doc.add_heading('Attendance and Punctuality Policy', level=1)
    doc.add_paragraph(
        'Effective Date: January 1, 2026\n'
        'Approved by: Human Resources Department'
    )
    doc.add_heading('1. Purpose', level=2)
    doc.add_paragraph(
        'This policy establishes expectations for employee attendance and '
        'punctuality at Meridian Solutions Inc. Regular and timely attendance '
        'is essential to maintaining operational efficiency and team collaboration.'
    )
    doc.add_heading('2. Work Hours', level=2)
    doc.add_paragraph(
        'Standard work hours are 9:00 AM to 5:30 PM, Monday through Friday. '
        'Core hours during which all employees must be available are 10:00 AM '
        'to 4:00 PM. Flexible start and end times may be arranged with manager approval.'
    )
    doc.add_heading('3. Reporting Absences', level=2)
    doc.add_paragraph(
        'Employees must notify their direct manager at least 30 minutes before '
        'their scheduled start time if they will be absent or late. Notification '
        'should be via phone or email. Repeated unexcused absences may result in '
        'disciplinary action up to and including termination.'
    )
    doc.add_heading('4. Tardiness', level=2)
    doc.add_paragraph(
        'Arriving more than 10 minutes after the scheduled start time is considered '
        'tardiness. Three instances of tardiness within a 30-day period will be '
        'addressed through the performance improvement process.'
    )
    doc.add_heading('5. Tracking', level=2)
    doc.add_paragraph(
        'All employees are required to record their attendance using the company '
        'HR system (HRConnect). Time records must be submitted and approved by '
        'the last working day of each pay period.'
    )
    doc.save(path)
    print(f'Created: {path}')


def create_leave_policy():
    path = f'{WORKDIR}/leave_policy.docx'
    doc = Document()
    doc.add_heading('Leave of Absence Policy', level=1)
    doc.add_paragraph(
        'Effective Date: January 1, 2026\n'
        'Approved by: Human Resources Department'
    )
    doc.add_heading('1. Overview', level=2)
    doc.add_paragraph(
        'Meridian Solutions Inc. provides employees with several types of leave '
        'to support work-life balance and personal well-being. This policy outlines '
        'the types of leave available and the procedures for requesting them.'
    )
    doc.add_heading('2. Annual Leave', level=2)
    doc.add_paragraph(
        'Full-time employees accrue 1.5 days of annual leave per month (18 days '
        'per year). Part-time employees accrue leave on a pro-rata basis. Unused '
        'leave up to 10 days may be carried forward to the next calendar year.'
    )
    doc.add_heading('3. Sick Leave', level=2)
    doc.add_paragraph(
        'Employees are entitled to 10 days of paid sick leave per calendar year. '
        'For absences longer than 3 consecutive days, a medical certificate from a '
        'registered physician is required. Sick leave does not carry over year to year.'
    )
    doc.add_heading('4. Parental Leave', level=2)
    doc.add_paragraph(
        'Primary caregivers are entitled to 16 weeks of paid parental leave. '
        'Secondary caregivers are entitled to 4 weeks of paid leave. Adoption '
        'leave follows the same entitlements as parental leave.'
    )
    doc.add_heading('5. Application Process', level=2)
    doc.add_paragraph(
        'Annual leave requests must be submitted at least 5 business days in advance '
        'via HRConnect. Emergency leave requests will be reviewed on a case-by-case '
        'basis. All leave must be approved by the direct manager prior to commencement.'
    )
    doc.save(path)
    print(f'Created: {path}')


def create_remote_work_policy():
    path = f'{WORKDIR}/remote_work_policy.docx'
    doc = Document()
    doc.add_heading('Remote Work and Flexible Work Arrangements Policy', level=1)
    doc.add_paragraph(
        'Effective Date: March 1, 2026\n'
        'Approved by: Human Resources Department'
    )
    doc.add_heading('1. Purpose', level=2)
    doc.add_paragraph(
        'This policy defines the framework for remote work and hybrid work '
        'arrangements at Meridian Solutions Inc. The company supports flexible '
        'work options where operationally feasible and in alignment with team requirements.'
    )
    doc.add_heading('2. Eligibility', level=2)
    doc.add_paragraph(
        'Remote work eligibility is determined by job role, performance record, '
        'and manager discretion. Employees must have completed at least 6 months '
        'of continuous service and hold a performance rating of "Meets Expectations" '
        'or higher in the most recent review cycle.'
    )
    doc.add_heading('3. Work Environment Requirements', level=2)
    doc.add_paragraph(
        'Remote employees must maintain a dedicated, ergonomically appropriate '
        'workspace with reliable internet connectivity (minimum 25 Mbps download). '
        'All data security and confidentiality requirements apply equally to remote '
        'and in-office environments.'
    )
    doc.add_heading('4. Hybrid Schedule', level=2)
    doc.add_paragraph(
        'The default hybrid schedule requires a minimum of 3 days per week in the '
        'office. Departments with specialized equipment or client-facing roles may '
        'require a higher in-office frequency. Remote days should not coincide with '
        'scheduled team meetings or mandatory training sessions.'
    )
    doc.add_heading('5. Equipment and Support', level=2)
    doc.add_paragraph(
        'The company will provide a laptop and VPN access for approved remote workers. '
        'Additional peripherals (monitor, keyboard) may be requested and are subject '
        'to budget approval. IT support is available during standard business hours.'
    )
    doc.save(path)
    print(f'Created: {path}')


def create_travel_expense_policy():
    path = f'{WORKDIR}/travel_expense_policy.docx'
    doc = Document()
    doc.add_heading('Business Travel and Expense Reimbursement Policy', level=1)
    doc.add_paragraph(
        'Effective Date: January 1, 2026\n'
        'Approved by: Finance Department & Human Resources'
    )
    doc.add_heading('1. Scope', level=2)
    doc.add_paragraph(
        'This policy applies to all Meridian Solutions Inc. employees who incur '
        'business-related travel and entertainment expenses. Expenses must be '
        'reasonable, necessary, and directly related to business activities.'
    )
    doc.add_heading('2. Pre-Approval Requirements', level=2)
    doc.add_paragraph(
        'All business travel must be pre-approved by the employee\'s direct manager '
        'and, for international travel, by the department VP. Travel requests must '
        'be submitted at least 10 business days in advance using the Concur system.'
    )
    doc.add_heading('3. Accommodation', level=2)
    doc.add_paragraph(
        'Hotel bookings must not exceed USD 220 per night for domestic travel '
        'and USD 300 per night for international travel. Employees should book '
        'through the company\'s preferred travel portal to access negotiated rates. '
        'Exceptions must be pre-approved by Finance.'
    )
    doc.add_heading('4. Meals and Per Diem', level=2)
    doc.add_paragraph(
        'Daily meal allowances: Breakfast USD 15, Lunch USD 25, Dinner USD 50. '
        'Client entertainment meals are reimbursed at actual cost with receipts '
        'and require attendee list and business purpose documentation.'
    )
    doc.add_heading('5. Submission and Reimbursement', level=2)
    doc.add_paragraph(
        'All expense reports must be submitted within 30 days of the expense date. '
        'Original receipts are required for any single expense over USD 25. '
        'Approved claims are processed within 10 business days via payroll.'
    )
    doc.save(path)
    print(f'Created: {path}')


def ensure_no_hr_folder():
    """Make sure the HR_Policies_2026 folder does NOT exist in the initial state."""
    folder = '/home/user/Desktop/HR_Policies_2026'
    if os.path.exists(folder):
        import shutil
        shutil.rmtree(folder)
        print(f'Removed pre-existing folder: {folder}')


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    # Ensure the target folder doesn't exist yet
    ensure_no_hr_folder()

    # Create 4 policy documents
    create_attendance_policy()
    create_leave_policy()
    create_remote_work_policy()
    create_travel_expense_policy()

    print('All 4 policy documents created on Desktop.')

    # GUI-ready startup: open the file manager showing the Desktop folder
    # so the agent can see the files and perform the task
    launch_gui('nautilus /home/user/Desktop', delay_sec=2.0)
    print('GUI_READY: launched Nautilus file manager with DISPLAY=:0')


create_initial()
