"""
Initial Setup: magazine_article.docx for writer_obj_072
Task ID: writer_obj_072
Domain: libreoffice_writer

Creates a magazine article document with a title and body text on page 1.
No text boxes are present. The agent task is to add four linked text boxes.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_072'
OUTPUT = f'{WORKDIR}/Desktop/magazine_article.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # Set page size to A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Title at the top of page 1
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run("The Future of Urban Living")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.space_after = Pt(18)
    sub_run = subtitle_para.add_run("How cities are transforming to meet tomorrow's challenges")
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.name = "Calibri"
    sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # Author and date line
    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_para.paragraph_format.space_after = Pt(24)
    meta_run = meta_para.add_run("By Alexandra Mercer  |  March 2025  |  Urban Studies Journal")
    meta_run.font.size = Pt(10)
    meta_run.font.name = "Calibri"
    meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Divider paragraph
    div_para = doc.add_paragraph()
    div_para.paragraph_format.space_after = Pt(18)

    # Body text paragraphs with realistic magazine article content
    body_paragraphs = [
        (
            "As metropolitan populations surge past record thresholds, urban planners, "
            "architects, and policy makers are racing to reimagine the very fabric of city life. "
            "From vertical forests in Milan to underground parks in New York, the innovations "
            "reshaping our cities are as bold as they are necessary."
        ),
        (
            "The concept of the '15-minute city' — popularised by Paris Mayor Anne Hidalgo and "
            "urban theorist Carlos Moreno — proposes that residents should be able to reach "
            "all their daily necessities within a quarter-hour walk or cycle. Barcelona's "
            "superblock programme has already demonstrated the model's viability, reclaiming "
            "street space from cars and returning it to community life."
        ),
        (
            "Technology plays an increasingly central role in this transformation. Smart sensors "
            "embedded in road surfaces monitor traffic flow and air quality in real time. "
            "Adaptive traffic signals reduce commute times by up to 25%, while integrated "
            "mobility apps allow commuters to seamlessly switch between buses, trains, "
            "bike-shares, and ride-hailing services with a single tap."
        ),
        (
            "Green infrastructure is no longer optional. The 2024 Urban Climate Resilience "
            "Report found that cities with at least 30% tree canopy cover experienced average "
            "temperatures up to 4°C lower than comparable urban areas without significant "
            "greenery. Singapore's 'City in a Garden' vision has become a benchmark, "
            "integrating biodiversity corridors alongside commercial skyscrapers."
        ),
        (
            "Housing affordability remains the most pressing challenge. In London, Sydney, and "
            "San Francisco, median home prices have reached more than twelve times the median "
            "annual income. Co-living developments, micro-apartments, and community land trusts "
            "offer partial solutions, but experts agree that systemic change in zoning law and "
            "public investment will be essential to make cities truly inclusive."
        ),
    ]

    for body_text in body_paragraphs:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.paragraph_format.space_after = Pt(10)
        para.paragraph_format.first_line_indent = Cm(0.5)
        run = para.add_run(body_text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
