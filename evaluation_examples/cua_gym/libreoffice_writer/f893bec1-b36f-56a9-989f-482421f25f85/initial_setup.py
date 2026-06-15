"""
Initial Setup: Set background color of text box on page 1 to light yellow (#FFFDE7)
Task ID: writer_obj_019
Domain: libreoffice_writer

Creates callout_doc.docx with a text box (8cm x 4cm) containing a tip paragraph.
The text box has NO background color (transparent) in the initial state.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'callout_doc'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_textbox_no_fill(doc):
    """
    Add a text box (8cm x 4cm) with a tip paragraph and NO background fill.
    Uses raw XML to embed a WPS text box drawing in the document body.
    """
    # Dimensions in EMU (1 cm = 914400 EMU / 2.54)
    # 8cm wide, 4cm tall
    cx = int(8 * 914400 / 2.54)   # 2895600 EMU
    cy = int(4 * 914400 / 2.54)   # 1447800 EMU

    # Build the textbox XML with NO fill (noFill)
    textbox_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
              xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
              xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
              xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
              xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
              xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="114300" distR="114300">
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="1" name="Text Box 1"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr txBx="1">
                <a:spLocks noChangeArrowheads="1"/>
              </wps:cNvSpPr>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{cx}" cy="{cy}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:noFill/>
                <a:ln>
                  <a:solidFill>
                    <a:srgbClr val="4472C4"/>
                  </a:solidFill>
                </a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent>
                  <w:p>
                    <w:pPr>
                      <w:pStyle w:val="Normal"/>
                    </w:pPr>
                    <w:r>
                      <w:rPr>
                        <w:b/>
                        <w:color w:val="1F3864"/>
                        <w:sz w:val="22"/>
                      </w:rPr>
                      <w:t>Tip:</w:t>
                    </w:r>
                    <w:r>
                      <w:rPr>
                        <w:color w:val="1F3864"/>
                        <w:sz w:val="22"/>
                      </w:rPr>
                      <w:t xml:space="preserve"> Always save your document regularly to avoid losing progress. Use Ctrl+S as a keyboard shortcut for quick saving.</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr insFmtShape="l" ins="91440" t="91440" r="91440" b="91440">
                <a:normAutofit/>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>'''

    textbox_elem = etree.fromstring(textbox_xml)
    doc.element.body.append(textbox_elem)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Add document title
    title = doc.add_heading('Project Planning Guide', level=1)

    # Add introductory paragraph
    intro = doc.add_paragraph(
        'This document provides essential guidelines for effective project planning and execution. '
        'Following these best practices will help ensure your projects are completed on time and '
        'within budget.'
    )

    # Add a section heading
    doc.add_heading('Getting Started', level=2)

    # Add content paragraph
    doc.add_paragraph(
        'Before beginning any project, it is important to clearly define the scope, objectives, '
        'and deliverables. Schedule an initial kickoff meeting with all stakeholders to align '
        'on expectations and timelines.'
    )

    # Add a second content paragraph
    doc.add_paragraph(
        'Create a project charter that outlines the goals, team members, resources required, '
        'and key milestones. This document will serve as a reference throughout the project lifecycle.'
    )

    # Add the text box with NO fill (transparent background)
    add_textbox_no_fill(doc)

    # Add more content below the text box
    doc.add_heading('Planning Phase', level=2)
    doc.add_paragraph(
        'Develop a detailed work breakdown structure (WBS) to identify all tasks and subtasks. '
        'Assign responsibilities to team members and establish realistic deadlines for each deliverable.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
