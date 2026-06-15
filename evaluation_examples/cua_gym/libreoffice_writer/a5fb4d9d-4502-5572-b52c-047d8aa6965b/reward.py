"""
Reward Script: Process two Jupyter notebooks, extract code cells to .py files,
               and create a combined comparison report in combined_report.odt
Task ID: osworld_multi_apps_code_to_writer_file_012
Domain: libreoffice_writer (multi-app: browser + python + writer)

Scoring Rubric:
  Component 1 (0.20): timeseries_code.py exists on Desktop with proper # Cell N comments
  Component 2 (0.20): query_code.py exists on Desktop with proper # Cell N comments
  Component 3 (0.25): combined_report.odt exists on Desktop with a comparison table (5 columns)
  Component 4 (0.175): Table row for timeseries notebook shows correct code cell count (43)
  Component 5 (0.175): Table row for query notebook shows correct code cell count (28)
  Total: 1.0
"""

import os
import re

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_code_to_writer_file_012'


def extract_all_text_from_doc(doc):
    """Extract all text content from a docx document."""
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return texts


def get_table_cell_text(table, row_idx, col_idx):
    """Get text from a table cell safely."""
    try:
        return table.rows[row_idx].cells[col_idx].text.strip()
    except (IndexError, AttributeError):
        return ""


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # ----------------------------------------------------------------
    # Component 1: timeseries_code.py exists with proper # Cell N comments (0.20 points)
    # This FAILS on initial_env (empty Desktop) and PASSES on golden_env
    # ----------------------------------------------------------------
    ts_path = os.path.join(WORKDIR, 'timeseries_code.py')
    ts_cells_found = 0
    try:
        if os.path.exists(ts_path):
            with open(ts_path, 'r') as f:
                ts_content = f.read()
            ts_lines = ts_content.split('\n')
            # Count # Cell N comments with global numbering
            cell_comments = [l for l in ts_lines if re.match(r'^# Cell \d+\s*$', l.strip())]
            ts_cells_found = len(cell_comments)
            if ts_cells_found >= 10:
                # At least 10 # Cell N comments means it's a proper extraction
                print(f"PASS: Component 1 — timeseries_code.py found with {ts_cells_found} '# Cell N' comments (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 1 — timeseries_code.py found but only {ts_cells_found} '# Cell N' comments (expected >= 10)")
        else:
            print(f"FAIL: Component 1 — timeseries_code.py not found at {ts_path}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ----------------------------------------------------------------
    # Component 2: query_code.py exists with proper # Cell N comments (0.20 points)
    # This FAILS on initial_env (empty Desktop) and PASSES on golden_env
    # ----------------------------------------------------------------
    query_path = os.path.join(WORKDIR, 'query_code.py')
    query_cells_found = 0
    try:
        if os.path.exists(query_path):
            with open(query_path, 'r') as f:
                query_content = f.read()
            query_lines = query_content.split('\n')
            # Count # Cell N comments
            cell_comments = [l for l in query_lines if re.match(r'^# Cell \d+\s*$', l.strip())]
            query_cells_found = len(cell_comments)
            if query_cells_found >= 5:
                # At least 5 # Cell N comments means it's a proper extraction
                print(f"PASS: Component 2 — query_code.py found with {query_cells_found} '# Cell N' comments (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — query_code.py found but only {query_cells_found} '# Cell N' comments (expected >= 5)")
        else:
            print(f"FAIL: Component 2 — query_code.py not found at {query_path}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ----------------------------------------------------------------
    # Component 3: combined_report.odt exists with a comparison table having 5 columns (0.25 points)
    # This FAILS on initial_env (empty Desktop) and PASSES on golden_env
    # ----------------------------------------------------------------
    odt_path = os.path.join(WORKDIR, 'combined_report.odt')
    doc = None
    try:
        if not os.path.exists(odt_path):
            print(f"FAIL: Component 3 — combined_report.odt not found at {odt_path}")
        else:
            # The file may be in .odt or .docx format (LibreOffice may save as docx)
            # Try python-docx first (handles OOXML)
            format_type = 'unknown'
            try:
                from docx import Document
                doc = Document(odt_path)
                format_type = 'docx/ooxml'
            except Exception as docx_err:
                doc = None

            if doc is None:
                print(f"FAIL: Component 3 — Cannot read combined_report.odt: not a supported format")
            else:
                # Check for a comparison table
                tables = doc.tables
                if len(tables) == 0:
                    print(f"FAIL: Component 3 — combined_report.odt has no tables (format: {format_type})")
                else:
                    # Find the main comparison table (should have 5 columns: Name, Total, Code, MD, LineCount)
                    main_table = None
                    for t in tables:
                        if len(t.columns) >= 4:
                            main_table = t
                            break
                    if main_table is None:
                        print(f"FAIL: Component 3 — No table with >= 4 columns found in combined_report.odt (tables: {len(tables)})")
                    else:
                        num_cols = len(main_table.columns)
                        num_rows = len(main_table.rows)
                        # Check header row contains key column names
                        header_texts = [main_table.rows[0].cells[c].text.strip().lower() for c in range(num_cols)]
                        has_notebook_col = any('notebook' in h or 'name' in h for h in header_texts)
                        has_code_col = any('code' in h for h in header_texts)
                        has_cell_col = any('cell' in h or 'total' in h or 'count' in h for h in header_texts)
                        if num_cols >= 4 and has_notebook_col and num_rows >= 3:
                            print(f"PASS: Component 3 — combined_report.odt has a comparison table ({num_rows} rows x {num_cols} cols) with notebook column (0.25 pts)")
                            total_score += 0.25
                        else:
                            print(f"FAIL: Component 3 — Table found but doesn't match requirements: {num_rows}x{num_cols}, headers={header_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ----------------------------------------------------------------
    # Component 4: Table row for timeseries notebook shows correct code cell count (0.175 points)
    # The task requires extracting code cells; the report should reflect this count correctly (43 from the notebook)
    # This FAILS on initial_env (no file) and PASSES on golden_env (correct data)
    # ----------------------------------------------------------------
    try:
        if doc is not None and len(doc.tables) > 0:
            main_table = None
            for t in doc.tables:
                if len(t.columns) >= 4:
                    main_table = t
                    break
            if main_table and len(main_table.rows) >= 2:
                # Search rows 1+ for timeseries notebook row
                ts_row = None
                for ri in range(1, len(main_table.rows)):
                    row_texts = [main_table.rows[ri].cells[c].text.strip() for c in range(len(main_table.columns))]
                    full_text = ' '.join(row_texts).lower()
                    if 'time' in full_text or '03.11' in full_text or 'timeseries' in full_text:
                        ts_row = row_texts
                        break
                if ts_row is None:
                    print(f"FAIL: Component 4 — No timeseries/03.11 row found in comparison table")
                else:
                    # Look for a numeric value >= 30 (code cell count for timeseries is 43)
                    # Allow some tolerance since the task says "approximate"
                    numeric_vals = []
                    for cell_text in ts_row:
                        try:
                            numeric_vals.append(int(cell_text))
                        except ValueError:
                            pass
                    # Code cell count should be in range [35, 55] (approximate)
                    code_cell_found = any(35 <= v <= 55 for v in numeric_vals)
                    if code_cell_found:
                        matching_val = [v for v in numeric_vals if 35 <= v <= 55]
                        print(f"PASS: Component 4 — Timeseries row has code cell count in valid range {matching_val} (expected ~43) (0.175 pts)")
                        total_score += 0.175
                    else:
                        print(f"FAIL: Component 4 — Timeseries row numeric values {numeric_vals} don't include expected code cell count (~43)")
        else:
            print(f"FAIL: Component 4 — No document or table available for timeseries row check")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ----------------------------------------------------------------
    # Component 5: Table row for query notebook shows correct code cell count (0.175 points)
    # The task requires extracting code cells; the report should reflect this count correctly (28 from the notebook)
    # This FAILS on initial_env (no file) and PASSES on golden_env (correct data)
    # ----------------------------------------------------------------
    try:
        if doc is not None and len(doc.tables) > 0:
            main_table = None
            for t in doc.tables:
                if len(t.columns) >= 4:
                    main_table = t
                    break
            if main_table and len(main_table.rows) >= 3:
                # Search rows 1+ for query notebook row
                query_row = None
                for ri in range(1, len(main_table.rows)):
                    row_texts = [main_table.rows[ri].cells[c].text.strip() for c in range(len(main_table.columns))]
                    full_text = ' '.join(row_texts).lower()
                    if 'query' in full_text or 'eval' in full_text or '03.12' in full_text or 'performance' in full_text:
                        query_row = row_texts
                        break
                if query_row is None:
                    print(f"FAIL: Component 5 — No query/eval/03.12 row found in comparison table")
                else:
                    # Look for a numeric value in range [20, 38] (code cell count for query is 28)
                    # Allow some tolerance since the task says "approximate"
                    numeric_vals = []
                    for cell_text in query_row:
                        try:
                            numeric_vals.append(int(cell_text))
                        except ValueError:
                            pass
                    # Code cell count should be in range [20, 38] (approximate)
                    code_cell_found = any(20 <= v <= 38 for v in numeric_vals)
                    if code_cell_found:
                        matching_val = [v for v in numeric_vals if 20 <= v <= 38]
                        print(f"PASS: Component 5 — Query row has code cell count in valid range {matching_val} (expected ~28) (0.175 pts)")
                        total_score += 0.175
                    else:
                        print(f"FAIL: Component 5 — Query row numeric values {numeric_vals} don't include expected code cell count (~28)")
        else:
            print(f"FAIL: Component 5 — No document or table available for query row check")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 3), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point — verify against Desktop
if __name__ == '__main__':
    verify_task()
else:
    verify_task()
