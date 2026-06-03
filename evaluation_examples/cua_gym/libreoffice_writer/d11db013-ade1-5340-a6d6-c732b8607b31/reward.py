"""
Reward Script: Enable header with bottom border and left-aligned text
Task ID: writer_page_074
Domain: libreoffice_writer
Scoring:
  - Component 1: Header contains text 'Operations Manual v3.2'     (0.4 pts)
  - Component 2: Header paragraph is left-aligned                  (0.3 pts)
  - Component 3: Header paragraph has a bottom border (single)     (0.3 pts)
Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_074'

EXPECTED_HEADER_TEXT = 'Operations Manual v3.2'


def get_header_para(doc):
    """Return the first header paragraph of the first section, or None."""
    if not doc.sections:
        return None
    header = doc.sections[0].header
    return header.paragraphs[0] if header.paragraphs else None


def check_header_text(para):
    """Return True if header paragraph text exactly matches expected."""
    if para is None:
        return False
    return para.text.strip() == EXPECTED_HEADER_TEXT


def check_header_alignment_left(para):
    """
    Return True if header paragraph has explicit left alignment in XML.
    Initial VM: no w:jc element (defaults, no explicit jc).
    Golden VM: w:jc w:val="left" present.
    """
    if para is None:
        return False
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return False
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        # Also check via python-docx alignment attr (LEFT == 0)
        fmt_align = para.paragraph_format.alignment
        return fmt_align is not None and int(fmt_align) == 0
    jc_val = jc.get(qn('w:val'))
    return jc_val in ('left', 'start')


def check_header_bottom_border(para):
    """
    Return (has_border, border_info_str) where has_border is True if
    header paragraph has a non-nil w:bottom border in w:pBdr.
    """
    if para is None:
        return False, "no paragraph"
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return False, "no pPr"
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        return False, "no pBdr"
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is None:
        return False, "no w:bottom in pBdr"
    border_val = bottom.get(qn('w:val'), '')
    if border_val and border_val not in ('none', 'nil'):
        border_color = bottom.get(qn('w:color'), 'auto')
        return True, f"val={border_val}, color={border_color}"
    return False, f"border val='{border_val}' is nil/none"


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Enable the header with a bottom border line separating it from the
    main content, and add the text 'Operations Manual v3.2' left-aligned.

    Ground truth (from context):
      - header enabled (has content)
      - header text 'Operations Manual v3.2' left-aligned
      - a bottom border line on the header paragraph (thin solid line)
      - footer unchanged (centered page numbers)
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section
    try:
        if len(doc.sections) < 1:
            print("CRITICAL: No sections found in document")
            print("REWARD: 0.0")
            return 0.0
    except Exception as e:
        print(f"CRITICAL: Cannot access document sections: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Header contains text 'Operations Manual v3.2' (0.4 points)
    # This FAILS on initial (empty header) and PASSES on golden
    try:
        para = get_header_para(doc)
        header_text = para.text.strip() if para else ''
        if check_header_text(para):
            print(f"PASS: Component 1 — header text matches '{EXPECTED_HEADER_TEXT}' (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — expected header text '{EXPECTED_HEADER_TEXT}', found '{header_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header paragraph is left-aligned (0.3 points)
    # In the golden file: w:jc w:val="left" present in pPr
    # In the initial file: no w:jc element (None alignment, defaults to left visually but not set)
    # We require explicit left alignment to distinguish task-introduced change
    try:
        para = get_header_para(doc)
        if check_header_alignment_left(para):
            print("PASS: Component 2 — header paragraph is explicitly left-aligned (0.3 pts)")
            total_score += 0.3
        else:
            fmt_align = para.paragraph_format.alignment if para else 'N/A'
            print(f"FAIL: Component 2 — header paragraph not explicitly left-aligned (alignment={fmt_align})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header paragraph has a bottom border (single line) (0.3 points)
    # In the golden file: <w:pBdr><w:bottom w:val="single" .../></w:pBdr>
    # In the initial file: no pBdr element at all
    try:
        para = get_header_para(doc)
        has_border, border_info = check_header_bottom_border(para)
        if has_border:
            print(f"PASS: Component 3 — header has bottom border ({border_info}) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — header paragraph has no bottom border ({border_info})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/ops_manual.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
