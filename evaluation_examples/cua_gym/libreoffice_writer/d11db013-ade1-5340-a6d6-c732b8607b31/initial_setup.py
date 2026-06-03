"""
Initial Setup: Operations Manual document without header
Task ID: writer_page_074
Domain: libreoffice_writer

Creates a 10-page operations manual (ops_manual.docx) with:
- A4 portrait, margins 2.54cm on all sides
- No header
- Footer with centered page numbers
- Realistic operations manual content spanning ~10 pages
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'ops_manual'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_to_footer(section):
    """Add centered page number field to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing paragraphs
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add "Page X" style field
    run_text = fp.add_run("Page ")
    # Begin field
    r1 = fp.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fldChar_begin)

    # Instruction text
    r2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    r2._r.append(instrText)

    # End field
    r3 = fp.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r3._r.append(fldChar_end)


def create_initial():
    doc = Document()

    # --- Page setup: A4 portrait, 2.54cm margins ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # --- No header: ensure header is empty/unlinked but contains no content ---
    header = section.header
    header.is_linked_to_previous = False
    # Leave header paragraphs empty (no text, no border)

    # --- Footer with centered page numbers ---
    add_page_number_to_footer(section)

    # ===== Page 1: Title Page =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Operations Manual')
    title_run.bold = True
    title_run.font.size = Pt(28)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('Version 3.2')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)

    doc.add_paragraph()

    dept_para = doc.add_paragraph()
    dept_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_para.add_run('Department of Operations & Logistics').font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.add_run('Issued: January 15, 2025').font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    confidential_para = doc.add_paragraph()
    confidential_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    conf_run = confidential_para.add_run('CONFIDENTIAL — Internal Use Only')
    conf_run.bold = True
    conf_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    conf_run.font.size = Pt(11)

    doc.add_page_break()

    # ===== Page 2: Table of Contents =====
    toc_heading = doc.add_heading('Table of Contents', level=1)

    toc_entries = [
        ('1. Introduction & Scope', '3'),
        ('2. Organizational Structure', '4'),
        ('3. Standard Operating Procedures', '5'),
        ('4. Safety & Compliance', '6'),
        ('5. Equipment & Asset Management', '7'),
        ('6. Incident Management', '8'),
        ('7. Vendor & Supplier Relations', '9'),
        ('8. Performance Metrics', '10'),
        ('9. Training & Development', '11'),
        ('10. Document Revision History', '12'),
    ]

    for entry, page_num in toc_entries:
        toc_para = doc.add_paragraph()
        toc_para.add_run(entry)
        toc_para.add_run('\t' + page_num)

    doc.add_page_break()

    # ===== Page 3: Introduction & Scope =====
    doc.add_heading('1. Introduction & Scope', level=1)

    doc.add_paragraph(
        'This Operations Manual (the "Manual") has been prepared by the Department of Operations & Logistics '
        'to provide comprehensive guidance for all operational staff, supervisors, and management personnel '
        'at GlobalTech Manufacturing Inc. The procedures described herein apply to all facilities in '
        'the Asia-Pacific region, including Singapore, Kuala Lumpur, and Jakarta operations centers.'
    )

    doc.add_paragraph(
        'The Manual is effective from January 15, 2025, and supersedes all prior operational guidelines '
        'and standard procedure documents issued before this date. All employees with operational '
        'responsibilities are expected to familiarize themselves with this document and comply with '
        'its provisions in full.'
    )

    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'The primary purpose of this Manual is to ensure consistent, efficient, and safe operations '
        'across all facilities. By standardizing workflows, we minimize errors, reduce downtime, '
        'and maintain compliance with regulatory requirements including ISO 9001:2015 and OHSAS 18001.'
    )

    doc.add_heading('1.2 Applicability', level=2)
    doc.add_paragraph(
        'This document applies to: (a) full-time and part-time operational staff; (b) contract workers '
        'employed for periods exceeding 30 days; (c) supervisors, team leads, and department managers; '
        '(d) any personnel performing tasks covered under Section 3 — Standard Operating Procedures.'
    )

    doc.add_page_break()

    # ===== Page 4: Organizational Structure =====
    doc.add_heading('2. Organizational Structure', level=1)

    doc.add_paragraph(
        'GlobalTech Manufacturing Inc. operates under a hierarchical organizational model with clear '
        'reporting lines from the Chief Operations Officer (COO) down to individual production line operators. '
        'The following describes the primary organizational units relevant to daily operations.'
    )

    doc.add_heading('2.1 Senior Leadership', level=2)
    leadership_data = [
        ('Chief Operations Officer', 'Dr. Rebecca Thornton', 'r.thornton@globaltech.com'),
        ('VP Operations — APAC', 'Mr. James Kwok', 'j.kwok@globaltech.com'),
        ('Director of Logistics', 'Ms. Priya Ramaswamy', 'p.ramaswamy@globaltech.com'),
        ('Head of Compliance', 'Mr. Daniel Müller', 'd.muller@globaltech.com'),
        ('Safety Manager', 'Ms. Yuki Tanaka', 'y.tanaka@globaltech.com'),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Contact'
    for role, name, email in leadership_data:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = name
        row_cells[2].text = email

    doc.add_paragraph()
    doc.add_heading('2.2 Department Heads', level=2)
    doc.add_paragraph(
        'Each operational department is led by a Department Head who reports directly to the VP Operations. '
        'Department Heads are responsible for resource allocation, performance monitoring, and compliance '
        'within their respective units. Weekly status reports are submitted every Monday by 09:00 SGT.'
    )

    doc.add_page_break()

    # ===== Page 5: Standard Operating Procedures =====
    doc.add_heading('3. Standard Operating Procedures', level=1)

    doc.add_paragraph(
        'Standard Operating Procedures (SOPs) are documented step-by-step instructions designed to '
        'achieve efficiency and quality in the performance of a defined operation. All SOPs are reviewed '
        'annually and updated as required to reflect process improvements and regulatory changes.'
    )

    doc.add_heading('3.1 Production Line Startup', level=2)
    startup_steps = [
        'Verify shift handover report from previous shift supervisor.',
        'Conduct equipment visual inspection using Form OPS-001.',
        'Check raw material inventory against production schedule (SAP module PP).',
        'Power on main control panels in sequence: MCP-A → MCP-B → MCP-C.',
        'Run diagnostics using the Automated Line Check (ALC) system — wait for GREEN status.',
        'Notify Quality Assurance team to position First-Off inspector.',
        'Begin production run and log start time in the Digital Shift Log (DSL).',
    ]
    for i, step in enumerate(startup_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    doc.add_heading('3.2 Production Line Shutdown', level=2)
    shutdown_steps = [
        'Complete current production batch and log actual output in DSL.',
        'Power down conveyors in reverse order: MCC → MCP-C → MCP-B → MCP-A.',
        'Conduct post-production equipment inspection using Form OPS-002.',
        'Secure all raw materials and work-in-progress items.',
        'Complete shift handover report and brief incoming supervisor.',
        'Submit end-of-shift production summary to Operations Manager.',
    ]
    for i, step in enumerate(shutdown_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    doc.add_page_break()

    # ===== Page 6: Safety & Compliance =====
    doc.add_heading('4. Safety & Compliance', level=1)

    doc.add_paragraph(
        'Safety is the cornerstone of all operations at GlobalTech Manufacturing Inc. Zero-tolerance policies '
        'are enforced for safety violations. Every employee is required to complete mandatory safety training '
        'before commencing any operational duties.'
    )

    doc.add_heading('4.1 Personal Protective Equipment (PPE)', level=2)
    doc.add_paragraph(
        'The following PPE is mandatory in all production areas: hard hat (ANSI Z89.1 Class E), '
        'safety glasses (ANSI Z87.1), steel-toed boots (ASTM F2413), high-visibility vest (ANSI/ISEA 107), '
        'and hearing protection in areas exceeding 85 dB(A). Supervisors are responsible for PPE compliance checks.'
    )

    doc.add_heading('4.2 Emergency Procedures', level=2)
    emergency_items = [
        'Fire: Activate nearest alarm, evacuate via marked exits, assemble at muster point EP-3.',
        'Chemical Spill: Isolate area, don SCBA, contact HAZMAT team via extension 4444.',
        'Medical Emergency: Call First Aid (ext. 5555), do not move injured person unless immediate danger.',
        'Power Failure: Activate emergency lighting, secure equipment, await clearance from Engineering.',
        'Security Threat: Lock down workstation, contact Security (ext. 6666), do not confront threat.',
    ]
    for item in emergency_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.3 Regulatory Compliance', level=2)
    doc.add_paragraph(
        'All operations must adhere to applicable local and international regulations. Current compliance '
        'frameworks include: Workplace Safety and Health Act (Singapore), Factory and Machinery Act (Malaysia), '
        'ISO 45001:2018 Occupational Health and Safety, and ISO 14001:2015 Environmental Management.'
    )

    doc.add_page_break()

    # ===== Page 7: Equipment & Asset Management =====
    doc.add_heading('5. Equipment & Asset Management', level=1)

    doc.add_paragraph(
        'Effective management of equipment and physical assets is critical to operational continuity. '
        'All assets with a value exceeding SGD 1,000 must be registered in the Asset Management System (AMS) '
        'and assigned a unique Asset Tag Number (ATN).'
    )

    doc.add_heading('5.1 Preventive Maintenance Schedule', level=2)
    doc.add_paragraph(
        'Preventive maintenance is scheduled based on manufacturer recommendations and operational risk '
        'assessments. The following maintenance frequencies apply:'
    )

    maint_table = doc.add_table(rows=1, cols=4)
    maint_table.style = 'Table Grid'
    maint_hdr = maint_table.rows[0].cells
    maint_hdr[0].text = 'Equipment Type'
    maint_hdr[1].text = 'Daily Check'
    maint_hdr[2].text = 'Monthly Service'
    maint_hdr[3].text = 'Annual Overhaul'
    maint_data = [
        ('CNC Machines', 'Lubrication, coolant', 'Filter replacement', 'Full calibration'),
        ('Conveyor Systems', 'Belt tension, guards', 'Drive unit service', 'Structural inspection'),
        ('Pneumatic Tools', 'Air pressure, seals', 'Valve inspection', 'Replacement of wear parts'),
        ('Electrical Panels', 'Indicator lights', 'Thermal imaging', 'Full electrical audit'),
        ('Forklifts', 'Fluid levels, tyres', 'Brake & steering', 'Certification renewal'),
    ]
    for row_data in maint_data:
        row_cells = maint_table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph()
    doc.add_heading('5.2 Asset Disposal', level=2)
    doc.add_paragraph(
        'Assets that are beyond economic repair or have reached end-of-life must be disposed of in '
        'accordance with local environmental regulations. A Disposal Request Form (DRF-07) must be '
        'submitted to the Finance and Compliance departments for approval prior to disposal.'
    )

    doc.add_page_break()

    # ===== Page 8: Incident Management =====
    doc.add_heading('6. Incident Management', level=1)

    doc.add_paragraph(
        'All incidents, near-misses, and hazardous conditions must be reported and investigated '
        'promptly. The goal of incident management is to prevent recurrence and continuously improve '
        'the safety culture within the organization.'
    )

    doc.add_heading('6.1 Incident Classification', level=2)
    incident_classes = [
        'Class 1 — Near Miss: No injury or damage; potential for harm existed.',
        'Class 2 — Minor Incident: First-aid treatment required; no lost time.',
        'Class 3 — Lost Time Incident (LTI): Employee unable to perform normal duties next day.',
        'Class 4 — Serious Bodily Harm: Hospitalization required; reportable to authority.',
        'Class 5 — Fatality: Death of employee or contractor; immediate regulatory notification.',
    ]
    for item in incident_classes:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.2 Reporting Timeline', level=2)
    doc.add_paragraph(
        'Class 1 and 2 incidents must be reported to the immediate supervisor within 2 hours and '
        'formally documented in the Incident Management System (IMS) within 24 hours. '
        'Class 3 incidents require notification to the Safety Manager within 1 hour and submission '
        'of a full investigation report within 5 working days. Class 4 and 5 incidents must be '
        'reported to regulatory authorities within 10 days in accordance with statutory requirements.'
    )

    doc.add_page_break()

    # ===== Page 9: Vendor & Supplier Relations =====
    doc.add_heading('7. Vendor & Supplier Relations', level=1)

    doc.add_paragraph(
        'GlobalTech Manufacturing maintains a pre-qualified vendor list managed by the Procurement '
        'department. All purchases exceeding SGD 5,000 must be made from approved vendors or through '
        'a formal tender process. Single-source purchases above SGD 20,000 require written approval '
        'from the COO.'
    )

    doc.add_heading('7.1 Approved Vendor Categories', level=2)
    vendor_categories = [
        'Raw Materials & Components: Grade A suppliers with ISO 9001 certification required.',
        'MRO Supplies: Preferred suppliers listed in the Procurement Portal (PP-2025).',
        'Logistics & Freight: Contracted carriers — FastFreight Asia, LogiFirst, TranzRoute.',
        'IT Equipment: Dell Premier Partner, HP Business Direct, Lenovo Gold Partner.',
        'Professional Services: Pre-qualified consultants and contractors on Panel Agreement.',
    ]
    for item in vendor_categories:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 Supplier Performance Review', level=2)
    doc.add_paragraph(
        'Supplier performance is evaluated quarterly using a balanced scorecard covering: '
        'on-time delivery (40%), quality defect rate (30%), responsiveness (20%), and '
        'sustainability practices (10%). Suppliers scoring below 60% for two consecutive quarters '
        'are placed on Probationary Watch and subject to a formal improvement plan.'
    )

    doc.add_page_break()

    # ===== Page 10: Performance Metrics & Revision History =====
    doc.add_heading('8. Performance Metrics', level=1)

    doc.add_paragraph(
        'Key Performance Indicators (KPIs) are reviewed monthly at the Operations Review Meeting (ORM). '
        'The following metrics are tracked at department and facility levels.'
    )

    kpi_table = doc.add_table(rows=1, cols=3)
    kpi_table.style = 'Table Grid'
    kpi_hdr = kpi_table.rows[0].cells
    kpi_hdr[0].text = 'KPI'
    kpi_hdr[1].text = 'Target'
    kpi_hdr[2].text = 'Measurement Frequency'
    kpi_data = [
        ('Overall Equipment Effectiveness (OEE)', '≥ 85%', 'Daily'),
        ('On-Time Delivery Rate', '≥ 97%', 'Weekly'),
        ('Defect Rate (ppm)', '≤ 250 ppm', 'Monthly'),
        ('Safety Incident Rate (per 1M hours)', '< 1.5', 'Monthly'),
        ('Employee Training Compliance', '100%', 'Quarterly'),
        ('Inventory Turnover Ratio', '≥ 8.0', 'Quarterly'),
    ]
    for row_data in kpi_data:
        row_cells = kpi_table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph()
    doc.add_heading('10. Document Revision History', level=1)

    rev_table = doc.add_table(rows=1, cols=4)
    rev_table.style = 'Table Grid'
    rev_hdr = rev_table.rows[0].cells
    rev_hdr[0].text = 'Version'
    rev_hdr[1].text = 'Date'
    rev_hdr[2].text = 'Author'
    rev_hdr[3].text = 'Summary of Changes'
    rev_data = [
        ('1.0', '2021-03-01', 'R. Thornton', 'Initial release'),
        ('2.0', '2022-09-12', 'J. Kwok', 'Added APAC facilities, updated SOPs'),
        ('3.0', '2023-11-05', 'P. Ramaswamy', 'Incorporated ISO 45001 requirements'),
        ('3.1', '2024-06-18', 'D. Müller', 'Minor corrections, updated KPI targets'),
        ('3.2', '2025-01-15', 'Y. Tanaka', 'Added Vendor section, updated emergency procedures'),
    ]
    for row_data in rev_data:
        row_cells = rev_table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    # Ensure Desktop dir exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
