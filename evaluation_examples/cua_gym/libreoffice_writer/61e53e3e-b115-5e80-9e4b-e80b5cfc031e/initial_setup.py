"""
Initial Setup: Insert a callout shape on page 1 pointing to the second paragraph
Task ID: writer_obj_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'review_doc'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Page setup: standard letter page
    section = doc.sections[0]
    from docx.shared import Inches
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Paragraph 1 - Title/Introduction
    p1 = doc.add_paragraph()
    run1 = p1.add_run("Project Overview and Strategic Objectives")
    run1.bold = True
    run1.font.size = Pt(14)

    # Paragraph 2 - body paragraph (target for callout)
    p2 = doc.add_paragraph(
        "The Q3 budget allocation requires a detailed assessment of departmental "
        "spending, including personnel costs, infrastructure investments, and "
        "operational expenses. Current projections indicate a 12% increase in "
        "technology expenditures compared to the previous fiscal year."
    )
    p2.paragraph_format.space_after = Pt(8)

    # Paragraph 3 - body paragraph
    p3 = doc.add_paragraph(
        "Marketing and sales teams have submitted revised forecasts reflecting "
        "shifting consumer trends and expanded distribution channels. The new "
        "partnerships with regional vendors are expected to generate an additional "
        "$1.4 million in revenue by end of Q4."
    )
    p3.paragraph_format.space_after = Pt(8)

    # Paragraph 4 - body paragraph
    p4 = doc.add_paragraph(
        "Human resources will finalize the updated compensation benchmarks by "
        "September 30th. All department heads are required to submit feedback on "
        "the proposed salary bands and benefits package revisions before the "
        "executive review scheduled for October 15th."
    )
    p4.paragraph_format.space_after = Pt(8)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
