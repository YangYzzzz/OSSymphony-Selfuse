"""
Initial Setup: Customer Retention Analysis document with findings paragraph (no highlighting)
Task ID: writer_txtfmt_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title: "Customer Retention Analysis Q4"
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Customer Retention Analysis Q4")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Calibri"

    # Introduction paragraph
    intro_para = doc.add_paragraph()
    intro_run = intro_para.add_run(
        "This report summarizes the findings from our comprehensive customer retention study "
        "conducted during Q4. The study examined behavioral patterns, feedback surveys, and "
        "transaction data from over 12,000 active customers across five regional markets. "
        "The insights presented here are intended to inform strategic decisions for the upcoming fiscal year."
    )
    intro_run.font.size = Pt(11)
    intro_run.font.name = "Calibri"

    # Findings paragraph — NO highlighting (task is to add yellow highlight)
    findings_para = doc.add_paragraph()
    findings_run = findings_para.add_run(
        "Our analysis reveals three primary drivers of customer churn: "
        "(1) pricing sensitivity among the 25-34 demographic, "
        "(2) lack of personalized communication, and "
        "(3) delayed response times exceeding 48 hours."
    )
    findings_run.font.size = Pt(11)
    findings_run.font.name = "Calibri"
    # Explicitly ensure no highlight is set
    findings_run.font.highlight_color = None

    # Additional context paragraph
    conclusion_para = doc.add_paragraph()
    conclusion_run = conclusion_para.add_run(
        "Addressing these root causes is projected to improve Q1 retention by 8-12%. "
        "The following sections provide detailed breakdowns by segment, region, and time period, "
        "with recommended action items for each identified driver."
    )
    conclusion_run.font.size = Pt(11)
    conclusion_run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup — open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
