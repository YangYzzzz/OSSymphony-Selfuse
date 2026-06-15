"""
Reward Script: Highlight the key findings paragraph in yellow.
Task ID: writer_txtfmt_008
Domain: libreoffice_writer
Scoring:
  Component 1 (0.7 pts): The findings paragraph starting with 'Our analysis reveals...'
                          has character highlight color = YELLOW on ALL its text-bearing runs.
  Component 2 (0.3 pts): [Sub-check only awarded if Component 1 passes]
                          Other paragraphs have NO highlight (no spurious highlighting).
  Total: 1.0 pts

Note: Text preservation is checked as a precondition gate (not scored) to avoid false positives.
"""

import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_008'

FINDINGS_START = 'Our analysis reveals three primary drivers of customer churn:'
FINDINGS_EXPECTED_TEXT = (
    'Our analysis reveals three primary drivers of customer churn: '
    '(1) pricing sensitivity among the 25-34 demographic, '
    '(2) lack of personalized communication, and '
    '(3) delayed response times exceeding 48 hours.'
)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the findings paragraph by index (use index to avoid object identity issues)
    findings_para_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(FINDINGS_START):
            findings_para_idx = idx
            break

    if findings_para_idx is None:
        print(f"CRITICAL: Findings paragraph starting with '{FINDINGS_START}' not found in document.")
        print("REWARD: 0.0")
        return 0.0

    findings_para = doc.paragraphs[findings_para_idx]

    # Precondition gate: Text content of findings paragraph must be intact
    actual_text = findings_para.text.strip()
    expected_text = FINDINGS_EXPECTED_TEXT.strip()
    if actual_text != expected_text:
        print(f"GATE FAIL: Findings paragraph text is corrupted or changed.")
        print(f"  Expected: {expected_text!r}")
        print(f"  Found:    {actual_text!r}")
        print("REWARD: 0.0")
        return 0.0
    else:
        print("GATE PASS: Findings paragraph text is intact (precondition).")

    # Component 1: The findings paragraph has YELLOW highlight on ALL its text-bearing runs (0.7 points)
    # This FAILS on initial_env (no highlight) and PASSES on golden_env (yellow highlight).
    try:
        runs_with_text = [run for run in findings_para.runs if run.text.strip()]
        if not runs_with_text:
            print("FAIL: Component 1 — Findings paragraph has no text runs.")
        else:
            all_yellow = all(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in runs_with_text)
            if all_yellow:
                print(f"PASS: Component 1 — All {len(runs_with_text)} run(s) in findings paragraph have YELLOW highlight. (0.7 pts)")
                total_score += 0.7
            else:
                # Report details about which runs failed
                for i, run in enumerate(runs_with_text):
                    hl = run.font.highlight_color
                    print(f"FAIL: Component 1 — Run {i} highlight={hl!r}, expected YELLOW. text={run.text[:60]!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: No other paragraphs have been accidentally highlighted (0.3 points)
    # Only awarded when Component 1 passes (yellow is present on findings para).
    # If no highlight was applied at all (initial state), Component 1 fails so we never reach here
    # with a passing total — but to be safe, this component is gated on Component 1 passing.
    if total_score >= 0.7:
        try:
            spurious_highlights = []
            for i, para in enumerate(doc.paragraphs):
                if i == findings_para_idx:
                    continue
                for j, run in enumerate(para.runs):
                    if run.font.highlight_color is not None:
                        spurious_highlights.append(
                            f"Para {i} Run {j}: highlight={run.font.highlight_color!r} text={run.text[:40]!r}"
                        )
            if not spurious_highlights:
                print("PASS: Component 2 — No spurious highlighting on other paragraphs. (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — {len(spurious_highlights)} unexpected highlighted run(s) in non-findings paragraphs:")
                for detail in spurious_highlights:
                    print(f"  {detail}")
        except Exception as e:
            print(f"ERROR: Component 2 — {e}")
    else:
        print("SKIP: Component 2 — skipped because Component 1 did not pass.")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
