"""
Reward Script: Insert page numbers at the bottom center of every page
Task ID: writer_legal_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): Footer contains PAGE field code
  Component 2 (0.3 pts): Footer paragraph is center-aligned
  Component 3 (0.2 pts): Complete field code structure (begin + end fldChar pair)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_004'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one section
    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Check ALL sections for footer page numbers
    # Task says "every page", so we check all sections
    has_page_field = False
    has_center_align = False
    has_fld_pair = False

    for sec_idx, section in enumerate(doc.sections):
        footer = section.footer

        # Component 1: Footer contains PAGE field code (0.5 points)
        try:
            for para in footer.paragraphs:
                instr_texts = para._element.findall('.//w:instrText', ns)
                for it in instr_texts:
                    if it.text and 'PAGE' in it.text.upper():
                        has_page_field = True
                        print(f"FOUND: PAGE field in section {sec_idx} footer, instrText='{it.text}'")
                        break
                if has_page_field:
                    break
        except Exception as e:
            print(f"ERROR: Checking PAGE field in section {sec_idx}: {e}")

        # Component 2: Footer paragraph is center-aligned (0.3 points)
        try:
            for para in footer.paragraphs:
                align = para.paragraph_format.alignment
                if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    has_center_align = True
                    print(f"FOUND: Center alignment in section {sec_idx} footer")
                    break
                # Also check XML directly for jc center
                jc_elements = para._element.findall('.//w:jc', ns)
                for jc in jc_elements:
                    val = jc.get(qn('w:val'))
                    if val == 'center':
                        has_center_align = True
                        print(f"FOUND: Center alignment (XML) in section {sec_idx} footer")
                        break
                if has_center_align:
                    break
        except Exception as e:
            print(f"ERROR: Checking alignment in section {sec_idx}: {e}")

        # Component 3: Complete fldChar pair (begin + end) (0.2 points)
        try:
            for para in footer.paragraphs:
                fld_chars = para._element.findall('.//w:fldChar', ns)
                fld_types = [fc.get(qn('w:fldCharType')) for fc in fld_chars]
                if 'begin' in fld_types and ('end' in fld_types or 'separate' in fld_types):
                    has_fld_pair = True
                    print(f"FOUND: Complete fldChar pair in section {sec_idx} footer: {fld_types}")
                    break
        except Exception as e:
            print(f"ERROR: Checking fldChar pair in section {sec_idx}: {e}")

    # Component 1: PAGE field code (0.5 points)
    if has_page_field:
        print(f"PASS: Component 1 -- PAGE field code found in footer (0.5 pts)")
        total_score += 0.5
    else:
        print(f"FAIL: Component 1 -- No PAGE field code found in any footer")

    # Component 2: Center alignment (0.3 points)
    if has_center_align:
        print(f"PASS: Component 2 -- Footer paragraph is center-aligned (0.3 pts)")
        total_score += 0.3
    else:
        print(f"FAIL: Component 2 -- Footer paragraph is not center-aligned")

    # Component 3: Complete field structure (0.2 points)
    if has_fld_pair:
        print(f"PASS: Component 3 -- Complete fldChar begin/end pair found (0.2 pts)")
        total_score += 0.2
    else:
        print(f"FAIL: Component 3 -- Incomplete or missing fldChar structure in footer")

    final_score = round(min(total_score, 1.0), 1)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
