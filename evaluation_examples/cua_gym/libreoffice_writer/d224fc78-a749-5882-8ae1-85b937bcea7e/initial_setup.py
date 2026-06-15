"""
Initial Setup: Insert page numbers at the bottom center of every page in appellate brief.
Task ID: writer_legal_004
Domain: libreoffice_writer

Creates a 15-page appellate brief with realistic legal content.
NO page numbers or footer content present.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_centered_heading(doc, text, level=1, bold=True, size=14):
    """Add a centered heading paragraph."""
    para = doc.add_heading(text, level=level)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in para.runs:
        run.bold = bold
        run.font.size = Pt(size)
    return para


def add_body_paragraph(doc, text, indent_first=True, space_after=6):
    """Add a body paragraph with standard legal formatting."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 2.0
    if indent_first:
        para.paragraph_format.first_line_indent = Inches(0.5)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return para


def add_section_heading(doc, text):
    """Add a section heading for the brief."""
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.underline = True
    return para


def create_initial():
    doc = Document()

    # Page setup - standard legal brief format
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Ensure NO footer content
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""

    # === COVER PAGE ===
    for _ in range(4):
        doc.add_paragraph()

    # Court title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("IN THE UNITED STATES COURT OF APPEALS")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("FOR THE NINTH CIRCUIT")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("Case No. 24-15782")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    doc.add_paragraph()

    # Parties
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("GREENFIELD TECHNOLOGIES, INC.,")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("Plaintiff-Appellant,")
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("v.")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("PACIFIC RIM INDUSTRIES, LLC,")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("Defendant-Appellee.")
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("OPENING BRIEF OF PLAINTIFF-APPELLANT")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("Appeal from the United States District Court\nfor the Northern District of California\nHonorable Maria L. Santos, District Judge\nCase No. 3:22-cv-04891-MLS")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Attorney info
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for line in [
        "KATHERINE A. MORRISON",
        "JAMES T. CHEN",
        "Morrison & Chen LLP",
        "555 Montgomery Street, Suite 2200",
        "San Francisco, California 94111",
        "Telephone: (415) 555-0187",
        "Facsimile: (415) 555-0188",
        "",
        "Attorneys for Plaintiff-Appellant",
        "Greenfield Technologies, Inc."
    ]:
        r = p.add_run(line + "\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        if line.startswith("KATHERINE") or line.startswith("JAMES"):
            r.bold = True

    # Page break to TABLE OF CONTENTS
    doc.add_page_break()

    # === TABLE OF CONTENTS ===
    add_section_heading(doc, "TABLE OF CONTENTS")

    toc_items = [
        ("TABLE OF AUTHORITIES", "iii"),
        ("JURISDICTIONAL STATEMENT", "1"),
        ("STATEMENT OF THE ISSUES", "2"),
        ("STATEMENT OF THE CASE", "3"),
        ("    A. Factual Background", "3"),
        ("    B. Procedural History", "5"),
        ("SUMMARY OF THE ARGUMENT", "6"),
        ("ARGUMENT", "7"),
        ("    I.  THE DISTRICT COURT ERRED IN GRANTING SUMMARY JUDGMENT ON THE BREACH OF CONTRACT CLAIM", "7"),
        ("        A. The Contract Unambiguously Required Performance by March 15, 2023", "7"),
        ("        B. Pacific Rim Failed to Deliver Conforming Goods", "9"),
        ("    II. THE DISTRICT COURT ABUSED ITS DISCRETION IN EXCLUDING KEY EXPERT TESTIMONY", "10"),
        ("        A. Dr. Henderson's Testimony Met the Daubert Standard", "10"),
        ("        B. The Exclusion Was Prejudicial", "11"),
        ("    III. THE DAMAGES AWARD SHOULD BE REVERSED", "12"),
        ("CONCLUSION", "13"),
        ("CERTIFICATE OF COMPLIANCE", "14"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r2 = p.add_run(f"  {page}")
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)

    # Page break to TABLE OF AUTHORITIES
    doc.add_page_break()

    # === TABLE OF AUTHORITIES ===
    add_section_heading(doc, "TABLE OF AUTHORITIES")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Cases")
    r.bold = True
    r.underline = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    cases = [
        "Anderson v. Liberty Lobby, Inc., 477 U.S. 242 (1986)",
        "Celotex Corp. v. Catrett, 477 U.S. 317 (1986)",
        "Daubert v. Merrell Dow Pharmaceuticals, Inc., 509 U.S. 579 (1993)",
        "General Electric Co. v. Joiner, 522 U.S. 136 (1997)",
        "Matsushita Elec. Indus. Co. v. Zenith Radio Corp., 475 U.S. 574 (1986)",
        "Pacific Gas & Elec. Co. v. G.W. Thomas Drayage & Rigging Co., 69 Cal.2d 33 (1968)",
        "Trident Center v. Connecticut General Life Ins. Co., 847 F.2d 564 (9th Cir. 1988)",
        "United States v. Hinkson, 585 F.3d 1247 (9th Cir. 2009)",
    ]
    for case in cases:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(case)
        r.italic = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Statutes and Rules")
    r.bold = True
    r.underline = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    statutes = [
        "28 U.S.C. \u00a7 1291",
        "28 U.S.C. \u00a7 1332",
        "Federal Rule of Civil Procedure 56",
        "Federal Rule of Evidence 702",
        "U.C.C. \u00a7 2-601",
        "U.C.C. \u00a7 2-711",
    ]
    for s in statutes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(s)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

    # Page break to JURISDICTIONAL STATEMENT
    doc.add_page_break()

    # === JURISDICTIONAL STATEMENT ===
    add_section_heading(doc, "JURISDICTIONAL STATEMENT")

    add_body_paragraph(doc,
        "The district court had subject matter jurisdiction over this action pursuant to "
        "28 U.S.C. \u00a7 1332, as the parties are citizens of different states and the amount "
        "in controversy exceeds $75,000, exclusive of interest and costs. Plaintiff-Appellant "
        "Greenfield Technologies, Inc. is a Delaware corporation with its principal place of "
        "business in San Francisco, California. Defendant-Appellee Pacific Rim Industries, LLC "
        "is an Oregon limited liability company with its principal place of business in Portland, Oregon."
    )

    add_body_paragraph(doc,
        "This Court has jurisdiction over this appeal pursuant to 28 U.S.C. \u00a7 1291. "
        "The district court entered final judgment on October 12, 2024, granting Defendant's "
        "motion for summary judgment and dismissing all claims with prejudice. Plaintiff filed "
        "a timely notice of appeal on November 8, 2024, within the thirty-day period prescribed "
        "by Federal Rule of Appellate Procedure 4(a)(1)(A)."
    )

    # === STATEMENT OF THE ISSUES ===
    add_section_heading(doc, "STATEMENT OF THE ISSUES")

    issues = [
        "Whether the district court erred in granting summary judgment on Greenfield's breach "
        "of contract claim when the contract's delivery deadline was unambiguous and Pacific Rim "
        "conceded it delivered non-conforming goods after the contractual deadline.",

        "Whether the district court abused its discretion in excluding the expert testimony of "
        "Dr. Robert Henderson regarding industry standards for semiconductor component quality "
        "control, where Dr. Henderson has over twenty-five years of experience in semiconductor "
        "manufacturing and his methodology was peer-reviewed and generally accepted in the field.",

        "Whether the damages award should be reversed where the district court failed to account "
        "for lost profits, cover damages under U.C.C. \u00a7 2-711, and consequential damages arising "
        "from Pacific Rim's breach."
    ]
    for i, issue in enumerate(issues, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 2.0
        r = p.add_run(f"{i}. {issue}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

    # Page break to STATEMENT OF THE CASE
    doc.add_page_break()

    # === STATEMENT OF THE CASE ===
    add_section_heading(doc, "STATEMENT OF THE CASE")

    # A. Factual Background
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("A. Factual Background")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    factual_paragraphs = [
        "Greenfield Technologies, Inc. (\"Greenfield\") is a leading developer of advanced sensor "
        "systems used in autonomous vehicle navigation. In January 2022, Greenfield entered into "
        "a supply agreement with Pacific Rim Industries, LLC (\"Pacific Rim\") for the procurement "
        "of custom semiconductor components essential to Greenfield's next-generation LiDAR sensor "
        "array, designated the GF-7000 series. (ER 45-67.)",

        "The Supply Agreement, executed on January 18, 2022, required Pacific Rim to deliver "
        "500,000 units of the custom-designed SR-400 semiconductor chip by March 15, 2023. "
        "Section 4.2 of the Agreement stated: \"Seller shall deliver all conforming goods to "
        "Buyer's designated facility no later than March 15, 2023. Time is of the essence with "
        "respect to Seller's delivery obligations.\" (ER 52, \u00a7 4.2.) The total contract price "
        "was $12.4 million. (ER 53, \u00a7 5.1.)",

        "Pacific Rim repeatedly assured Greenfield that production was on schedule. In email "
        "correspondence dated September 2022, Pacific Rim's Vice President of Operations, "
        "Thomas Watanabe, confirmed that \"all manufacturing milestones are being met and we "
        "remain on track for the March delivery date.\" (ER 112.) Similar assurances were "
        "provided in December 2022 and February 2023. (ER 118-125.)",

        "Despite these assurances, Pacific Rim did not deliver any components by the March 15, "
        "2023 deadline. The first shipment of 200,000 units did not arrive until May 8, 2023, "
        "nearly two months late. (ER 130.) Moreover, independent testing by Greenfield's quality "
        "assurance team revealed that approximately 23% of the delivered chips failed to meet the "
        "specifications set forth in Exhibit A to the Agreement. (ER 145-152.) Specifically, the "
        "chips exhibited unacceptable levels of signal noise and thermal instability at operating "
        "temperatures above 85\u00b0C, rendering them unsuitable for automotive applications. (ER 148.)",

        "As a direct result of Pacific Rim's failure to deliver conforming goods on time, "
        "Greenfield was unable to meet its own contractual obligations to its primary customer, "
        "Aurora Mobility Systems, for the delivery of the GF-7000 sensor units. Aurora subsequently "
        "terminated its $34.5 million contract with Greenfield and sourced alternative components "
        "from a competitor. (ER 160-168.) Greenfield also incurred $2.8 million in costs to secure "
        "replacement semiconductor components from an alternative supplier at higher prices. (ER 172.)",
    ]
    for text in factual_paragraphs:
        add_body_paragraph(doc, text)

    # Page break
    doc.add_page_break()

    # B. Procedural History
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("B. Procedural History")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    procedural_paragraphs = [
        "Greenfield commenced this action on August 15, 2022, in the United States District Court "
        "for the Northern District of California, asserting claims for breach of contract, breach "
        "of the implied warranty of merchantability under U.C.C. \u00a7 2-314, and negligent "
        "misrepresentation. (ER 1-28.)",

        "Following extensive discovery, Pacific Rim moved for summary judgment on all claims on "
        "June 3, 2024. Greenfield opposed the motion and submitted, among other evidence, the "
        "expert report and declaration of Dr. Robert Henderson, a semiconductor industry expert "
        "with over twenty-five years of experience. Dr. Henderson opined that the delivered chips "
        "did not conform to industry standards for automotive-grade semiconductor components and "
        "that Pacific Rim's quality control processes fell below accepted industry norms. (ER 200-235.)",

        "The district court granted summary judgment in favor of Pacific Rim on October 12, 2024. "
        "The court held that the contract's delivery deadline was ambiguous because Section 7.3 "
        "contained a force majeure provision that \"could be read to excuse late performance under "
        "certain circumstances.\" (ER 280.) The court also excluded Dr. Henderson's testimony under "
        "Daubert, finding that his methodology was \"insufficiently rigorous\" and \"not adequately "
        "tied to the facts of this case.\" (ER 285.) Finally, the court found Greenfield's damages "
        "theory \"too speculative\" because it relied on projected future revenues from the Aurora "
        "contract. (ER 290.)",
    ]
    for text in procedural_paragraphs:
        add_body_paragraph(doc, text)

    # Page break to SUMMARY OF THE ARGUMENT
    doc.add_page_break()

    # === SUMMARY OF THE ARGUMENT ===
    add_section_heading(doc, "SUMMARY OF THE ARGUMENT")

    summary_paragraphs = [
        "The district court committed reversible error in three respects. First, the court erred "
        "in finding the contract's delivery deadline ambiguous. Section 4.2 plainly states that "
        "delivery must occur \"no later than March 15, 2023\" and that \"[t]ime is of the essence.\" "
        "The force majeure clause in Section 7.3 does not create ambiguity; it merely provides a "
        "limited excuse for non-performance due to enumerated extraordinary events, none of which "
        "Pacific Rim invoked or established.",

        "Second, the district court abused its discretion in excluding Dr. Henderson's expert "
        "testimony. Dr. Henderson is one of the foremost experts in semiconductor manufacturing "
        "quality control. His methodology\u2014comparative analysis of chip performance data against "
        "industry benchmarks\u2014is widely used and has been accepted by multiple federal courts. The "
        "district court's conclusory finding that his approach was \"insufficiently rigorous\" cannot "
        "withstand scrutiny under Daubert and its progeny.",

        "Third, the district court erred in characterizing Greenfield's damages as speculative. "
        "Greenfield presented a fully executed contract with Aurora Mobility Systems worth $34.5 "
        "million, concrete evidence of cover costs totaling $2.8 million, and detailed financial "
        "projections prepared by its CFO. These are precisely the types of damages recognized under "
        "U.C.C. \u00a7\u00a7 2-711 through 2-715. The district court's rejection of this evidence was "
        "legal error.",
    ]
    for text in summary_paragraphs:
        add_body_paragraph(doc, text)

    # Page break to ARGUMENT
    doc.add_page_break()

    # === ARGUMENT ===
    add_section_heading(doc, "ARGUMENT")

    # I. BREACH OF CONTRACT
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("I. THE DISTRICT COURT ERRED IN GRANTING SUMMARY JUDGMENT ON THE BREACH OF CONTRACT CLAIM")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    add_body_paragraph(doc,
        "This Court reviews de novo a district court's grant of summary judgment. Anderson v. "
        "Liberty Lobby, Inc., 477 U.S. 242, 248 (1986). Summary judgment is appropriate only "
        "when \"there is no genuine dispute as to any material fact and the movant is entitled "
        "to judgment as a matter of law.\" Fed. R. Civ. P. 56(a). In reviewing the record, this "
        "Court must view the evidence in the light most favorable to the nonmoving party. "
        "Matsushita Elec. Indus. Co. v. Zenith Radio Corp., 475 U.S. 574, 587 (1986)."
    )

    # A. Contract Unambiguous
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run("A. The Contract Unambiguously Required Performance by March 15, 2023")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    contract_paragraphs = [
        "The interpretation of a contract is a question of law reviewed de novo. Trident Center "
        "v. Connecticut General Life Ins. Co., 847 F.2d 564, 569 (9th Cir. 1988). A contract "
        "term is ambiguous only if it is \"reasonably susceptible to more than one interpretation.\" "
        "Pacific Gas & Elec. Co. v. G.W. Thomas Drayage & Rigging Co., 69 Cal.2d 33, 37 (1968).",

        "Section 4.2 of the Supply Agreement is a model of clarity. It states, in full: \"Seller "
        "shall deliver all conforming goods to Buyer's designated facility no later than March 15, "
        "2023. Time is of the essence with respect to Seller's delivery obligations. Any failure "
        "to deliver conforming goods by the specified date shall constitute a material breach of "
        "this Agreement.\" (ER 52.) There is nothing ambiguous about this language. The deadline "
        "is fixed; the consequences of non-compliance are expressly stated.",

        "The district court's reliance on Section 7.3, the force majeure clause, to find ambiguity "
        "was legal error. Section 7.3 provides: \"Neither party shall be liable for failure or delay "
        "in performance to the extent caused by circumstances beyond its reasonable control, "
        "including but not limited to acts of God, war, terrorism, government action, epidemic, "
        "fire, flood, or severe weather.\" (ER 55.) Critically, Pacific Rim never invoked the "
        "force majeure clause. It never identified any qualifying event. It never provided the "
        "written notice required by Section 7.3(b). The mere existence of a force majeure clause "
        "does not transform an otherwise clear deadline into an ambiguous one.",

        "Moreover, Pacific Rim's own course of conduct confirms the deadline was understood to be "
        "firm. Vice President Watanabe's repeated assurances that production was \"on track\" for "
        "the March delivery date demonstrate that Pacific Rim itself understood March 15, 2023 as "
        "a fixed, binding deadline\u2014not a flexible target subject to unilateral adjustment. (ER 112, "
        "118-125.)",
    ]
    for text in contract_paragraphs:
        add_body_paragraph(doc, text)

    # Page break
    doc.add_page_break()

    # B. Non-conforming goods
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run("B. Pacific Rim Failed to Deliver Conforming Goods")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    nonconforming_paragraphs = [
        "Even setting aside the late delivery, the goods Pacific Rim eventually delivered were "
        "non-conforming. Under U.C.C. \u00a7 2-601, a buyer may reject goods that \"fail in any respect "
        "to conform to the contract.\" The undisputed record shows that 23% of the delivered SR-400 "
        "chips failed to meet the specifications in Exhibit A to the Agreement. (ER 145-152.)",

        "The evidence of non-conformity is overwhelming. Greenfield's quality assurance team "
        "conducted standardized testing protocols\u2014the same protocols used throughout the semiconductor "
        "industry\u2014and documented signal noise levels exceeding the contractual maximum by 15-40% and "
        "thermal instability at temperatures well within the specified operating range. (ER 148-150.) "
        "Pacific Rim has never disputed these test results.",
    ]
    for text in nonconforming_paragraphs:
        add_body_paragraph(doc, text)

    # II. EXPERT TESTIMONY
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("II. THE DISTRICT COURT ABUSED ITS DISCRETION IN EXCLUDING KEY EXPERT TESTIMONY")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    add_body_paragraph(doc,
        "This Court reviews the exclusion of expert testimony for abuse of discretion. General "
        "Electric Co. v. Joiner, 522 U.S. 136, 141 (1997). A district court abuses its discretion "
        "when it makes an error of law or reaches a conclusion that is \"illogical, implausible, or "
        "without support in inferences that may be drawn from the record.\" United States v. Hinkson, "
        "585 F.3d 1247, 1251 (9th Cir. 2009)."
    )

    # A. Daubert standard
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run("A. Dr. Henderson's Testimony Met the Daubert Standard")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    daubert_paragraphs = [
        "Under Daubert v. Merrell Dow Pharmaceuticals, Inc., 509 U.S. 579 (1993), expert "
        "testimony is admissible if it is based on sufficient facts or data, is the product of "
        "reliable principles and methods, and the expert has reliably applied those principles "
        "to the facts of the case. Fed. R. Evid. 702.",

        "Dr. Robert Henderson holds a Ph.D. in Electrical Engineering from MIT and has over "
        "twenty-five years of experience in semiconductor manufacturing and quality control. He "
        "has published thirty-seven peer-reviewed articles on chip performance metrics and has "
        "testified as an expert in fourteen prior federal proceedings. (ER 200-205.) His "
        "methodology\u2014comparative analysis of actual chip performance data against established "
        "industry benchmarks published by the Semiconductor Industry Association\u2014is the standard "
        "approach used by quality control professionals throughout the industry. (ER 210-215.)",

        "The district court's cursory conclusion that Dr. Henderson's methodology was "
        "\"insufficiently rigorous\" ignores the substantial record demonstrating otherwise. The "
        "court failed to identify any specific deficiency in Dr. Henderson's analysis. It did not "
        "explain why comparative benchmarking\u2014a methodology accepted by courts nationwide\u2014was "
        "unreliable in this context. This is precisely the type of conclusory gatekeeping that "
        "constitutes an abuse of discretion.",
    ]
    for text in daubert_paragraphs:
        add_body_paragraph(doc, text)

    # Page break
    doc.add_page_break()

    # B. Prejudicial exclusion
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run("B. The Exclusion Was Prejudicial")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    prejudicial_paragraphs = [
        "The exclusion of Dr. Henderson's testimony was not harmless error. His testimony was "
        "central to establishing that the delivered chips were non-conforming and that Pacific "
        "Rim's quality control processes fell below industry standards. Without his testimony, "
        "Greenfield was left to rely solely on its own internal testing data, which the district "
        "court characterized as \"self-serving.\" (ER 287.)",

        "Had Dr. Henderson been permitted to testify, the jury would have heard from an independent, "
        "highly credentialed expert that the SR-400 chips objectively failed to meet both the "
        "contractual specifications and industry-wide quality standards. This testimony would have "
        "directly supported Greenfield's breach of contract claim and undermined Pacific Rim's "
        "defense that the goods were substantially conforming. The exclusion therefore prejudiced "
        "Greenfield's case on the merits.",
    ]
    for text in prejudicial_paragraphs:
        add_body_paragraph(doc, text)

    # III. DAMAGES
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("III. THE DAMAGES AWARD SHOULD BE REVERSED")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    damages_paragraphs = [
        "The district court's characterization of Greenfield's damages as \"too speculative\" "
        "cannot be sustained. Greenfield presented three categories of damages, each supported "
        "by concrete, admissible evidence.",

        "First, Greenfield is entitled to cover damages under U.C.C. \u00a7 2-712. Greenfield "
        "purchased replacement semiconductor components from Quantera Semiconductor at a cost "
        "of $15.2 million\u2014$2.8 million more than the contract price with Pacific Rim. The "
        "replacement purchase was made in good faith, without unreasonable delay, and for "
        "substantially similar components. (ER 172-178.) These are textbook cover damages.",

        "Second, Greenfield is entitled to consequential damages under U.C.C. \u00a7 2-715(2). "
        "The Aurora Mobility Systems contract was worth $34.5 million. Pacific Rim knew at the "
        "time of contracting that Greenfield needed the SR-400 chips to fulfill its obligations "
        "to Aurora. (ER 60, \u00a7 2.3 [recital stating purpose of purchase].) The loss of the Aurora "
        "contract was a foreseeable consequence of Pacific Rim's breach, and the damages are "
        "calculable with reasonable certainty based on the executed contract. (ER 160-168.)",

        "Third, Greenfield is entitled to incidental damages including storage costs for the "
        "non-conforming goods ($45,000), costs of inspection and testing ($128,000), and "
        "administrative expenses related to securing replacement components ($67,000). (ER 180-185.) "
        "These damages are well-documented and undisputed.",
    ]
    for text in damages_paragraphs:
        add_body_paragraph(doc, text)

    # Page break to CONCLUSION
    doc.add_page_break()

    # === CONCLUSION ===
    add_section_heading(doc, "CONCLUSION")

    add_body_paragraph(doc,
        "For the foregoing reasons, Plaintiff-Appellant Greenfield Technologies, Inc. "
        "respectfully requests that this Court reverse the district court's grant of summary "
        "judgment, reverse the exclusion of Dr. Henderson's expert testimony, reverse the "
        "damages determination, and remand this case for trial on the merits."
    )

    doc.add_paragraph()

    # Signature block
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.left_indent = Inches(3)
    r = p.add_run("Respectfully submitted,\n\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.left_indent = Inches(3)
    for line in [
        "_____________________________",
        "KATHERINE A. MORRISON",
        "JAMES T. CHEN",
        "Morrison & Chen LLP",
        "555 Montgomery Street, Suite 2200",
        "San Francisco, California 94111",
        "Telephone: (415) 555-0187",
        "",
        "Attorneys for Plaintiff-Appellant",
        "Greenfield Technologies, Inc.",
    ]:
        r = p.add_run(line + "\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        if "KATHERINE" in line or "JAMES" in line:
            r.bold = True

    p = doc.add_paragraph()
    r = p.add_run("Dated: November 22, 2024")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    # Page break to CERTIFICATE OF COMPLIANCE
    doc.add_page_break()

    # === CERTIFICATE OF COMPLIANCE ===
    add_section_heading(doc, "CERTIFICATE OF COMPLIANCE")

    add_body_paragraph(doc,
        "Pursuant to Federal Rule of Appellate Procedure 32(g), I certify that this brief "
        "complies with the type-volume limitation of Rule 32(a)(7)(B) because it contains "
        "12,847 words, excluding the parts of the brief exempted by Rule 32(f)."
    )

    add_body_paragraph(doc,
        "This brief complies with the typeface requirements of Rule 32(a)(5) and the type-style "
        "requirements of Rule 32(a)(6) because it has been prepared in a proportionally spaced "
        "typeface using Microsoft Word in 14-point Times New Roman font."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.left_indent = Inches(3)
    r = p.add_run("_____________________________\nKATHERINE A. MORRISON")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    # Additional filler to ensure we reach ~15 pages with double spacing
    doc.add_page_break()

    # === CERTIFICATE OF SERVICE ===
    add_section_heading(doc, "CERTIFICATE OF SERVICE")

    add_body_paragraph(doc,
        "I hereby certify that on November 22, 2024, I electronically filed the foregoing "
        "Opening Brief of Plaintiff-Appellant with the Clerk of the Court for the United States "
        "Court of Appeals for the Ninth Circuit by using the appellate CM/ECF system."
    )

    add_body_paragraph(doc,
        "I certify that all participants in the case are registered CM/ECF users and that service "
        "will be accomplished by the appellate CM/ECF system. The following counsel of record were "
        "served electronically:"
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.5)
    for line in [
        "David R. Nakamura",
        "Sarah E. Blackwell",
        "Nakamura & Blackwell, P.C.",
        "1200 SW Fifth Avenue, Suite 3400",
        "Portland, Oregon 97204",
        "Telephone: (503) 555-0294",
        "Email: dnakamura@nblaw.com",
        "Email: sblackwell@nblaw.com",
        "",
        "Attorneys for Defendant-Appellee",
        "Pacific Rim Industries, LLC",
    ]:
        r = p.add_run(line + "\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.left_indent = Inches(3)
    r = p.add_run("_____________________________\nKATHERINE A. MORRISON")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
