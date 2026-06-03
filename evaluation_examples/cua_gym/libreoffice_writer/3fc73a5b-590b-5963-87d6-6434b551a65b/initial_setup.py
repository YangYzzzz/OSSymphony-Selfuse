"""
Initial Setup: Create a Writer document with six plain installation steps
Task ID: writer_bs_056
Domain: libreoffice_writer

The document contains a title, intro text, "Installation Procedure" heading,
and six paragraphs in Default Paragraph Style (no numbering, no list style).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Software Installation Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introductory paragraphs ---
    intro1 = doc.add_paragraph(
        'This guide provides step-by-step instructions for installing '
        'DataVault Pro 4.2 on your workstation. Please ensure you have '
        'administrator privileges before proceeding.'
    )

    intro2 = doc.add_paragraph(
        'System requirements: Windows 10 or later, 8 GB RAM minimum, '
        '2 GB free disk space, .NET Framework 4.8 or higher. A stable '
        'internet connection is recommended for license activation.'
    )

    # --- Installation Procedure heading ---
    doc.add_heading('Installation Procedure', level=1)

    # --- Six plain paragraphs (Default Paragraph Style, NO numbering) ---
    steps = [
        'Download the installer package from the official DataVault Pro website '
        'by navigating to the Downloads section and selecting the version '
        'compatible with your operating system.',

        'Run setup.exe as an administrator by right-clicking the downloaded '
        'file and selecting "Run as administrator" from the context menu.',

        'Accept the End User License Agreement after carefully reviewing the '
        'terms and conditions presented in the setup wizard dialog.',

        'Choose the installation directory or accept the default path at '
        'C:\\Program Files\\DataVault Pro\\4.2 and confirm sufficient disk space '
        'is available.',

        'Configure the database connection settings by entering the server '
        'hostname, port number (default 5432), and authentication credentials '
        'provided by your IT department.',

        'Complete the installation and restart the computer when prompted to '
        'finalize the system registry entries and environment variable updates.',
    ]

    for step_text in steps:
        doc.add_paragraph(step_text)

    # --- Additional section after the steps ---
    doc.add_heading('Post-Installation Verification', level=1)
    doc.add_paragraph(
        'After restarting, launch DataVault Pro from the Start menu. The '
        'application should display the activation wizard on first run. Enter '
        'your license key to complete the setup process.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
