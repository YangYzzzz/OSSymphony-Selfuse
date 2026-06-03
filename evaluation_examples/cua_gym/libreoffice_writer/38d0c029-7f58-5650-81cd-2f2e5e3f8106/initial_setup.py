"""
Initial Setup: Product brochure with single-column layout
Task ID: writer_rd_084
Domain: libreoffice_writer

Creates a Writer document with:
- Introduction paragraph
- 9 feature descriptions
- Conclusion paragraph
All in single-column default layout.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_084'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('NovaTech Pro X500 — Product Brochure', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introduction ---
    intro_heading = doc.add_heading('Introduction', level=2)

    intro_text = (
        "The NovaTech Pro X500 represents a breakthrough in portable computing, "
        "combining enterprise-grade performance with an ultra-slim design that weighs "
        "just 1.2 kg. Engineered for professionals who demand reliability in every "
        "environment, the X500 has been tested across 47 countries and rated #1 in "
        "customer satisfaction by TechReview Global for three consecutive years. "
        "Whether you are presenting to a boardroom in Tokyo, coding at a startup in "
        "Berlin, or analyzing field data in rural Kenya, the X500 adapts seamlessly "
        "to your workflow."
    )
    p_intro = doc.add_paragraph(intro_text)
    p_intro.paragraph_format.space_after = Pt(12)

    # --- Features Section (9 features, all single-column) ---
    features_heading = doc.add_heading('Features', level=2)

    features = [
        (
            "Quantum-Core M3 Processor",
            "Powered by the latest Quantum-Core M3 chipset with 16 efficiency cores "
            "and 8 performance cores, delivering up to 4.8 GHz burst frequency. "
            "Benchmarks show a 62% improvement over the previous generation in "
            "multi-threaded workloads."
        ),
        (
            "NanoEdge 4K OLED Display",
            "A stunning 15.6-inch 4K OLED panel with 100% DCI-P3 color gamut, "
            "600 nits peak brightness, and a 120 Hz adaptive refresh rate. The micro-bezel "
            "design achieves a 94% screen-to-body ratio for immersive viewing."
        ),
        (
            "HyperFlow Cooling System",
            "Dual vapor-chamber heat pipes with graphene-enhanced thermal interface "
            "material keep surface temperatures below 35 degrees Celsius even under "
            "sustained load. Fan noise stays under 28 dB in standard mode."
        ),
        (
            "64 GB LPDDR5X Memory",
            "Ultra-fast 7500 MT/s memory bandwidth enables seamless multitasking "
            "across virtual machines, 3D rendering suites, and large language model "
            "inference workloads without throttling."
        ),
        (
            "2 TB PCIe Gen 5 SSD",
            "Sequential read speeds of 14,000 MB/s and write speeds of 12,000 MB/s "
            "ensure that boot times stay under 3 seconds and project compilation "
            "completes 40% faster than NVMe Gen 4 drives."
        ),
        (
            "Wi-Fi 7 & 5G Connectivity",
            "Tri-band Wi-Fi 7 with 320 MHz channels delivers up to 5.8 Gbps throughput. "
            "An integrated Qualcomm X75 5G modem provides sub-6 GHz and mmWave "
            "coverage for reliable connectivity anywhere."
        ),
        (
            "All-Day 100 Wh Battery",
            "A high-density lithium-silicon battery provides up to 18 hours of mixed "
            "use on a single charge. The 140 W GaN charger reaches 50% capacity in "
            "just 25 minutes via USB-C Power Delivery 3.1."
        ),
        (
            "MIL-STD-810H Durability",
            "Tested to withstand drops from 1.22 meters, extreme temperatures from "
            "-29 to 63 degrees Celsius, and humidity levels up to 95%. The magnesium-alloy "
            "chassis resists flex and torsion during transport."
        ),
        (
            "Studio-Grade Audio Suite",
            "Quad Harman Kardon speakers with Dolby Atmos spatial audio deliver room-filling "
            "sound. A four-microphone array with AI noise cancellation ensures crystal-clear "
            "conference calls even in noisy environments."
        ),
    ]

    for feat_title, feat_desc in features:
        p = doc.add_paragraph()
        run_title = p.add_run(feat_title + ': ')
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_desc = p.add_run(feat_desc)
        run_desc.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    # --- Conclusion ---
    conclusion_heading = doc.add_heading('Conclusion', level=2)

    conclusion_text = (
        "The NovaTech Pro X500 is available starting Q3 2025 at an MSRP of $2,499 "
        "for the base configuration. Enterprise volume pricing, extended warranty "
        "packages, and custom configuration options are available through our global "
        "partner network spanning 120 countries. Contact your NovaTech account manager "
        "or visit novatech.example.com/pro-x500 for detailed specifications, customer "
        "testimonials, and ordering information."
    )
    p_conclusion = doc.add_paragraph(conclusion_text)
    p_conclusion.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
