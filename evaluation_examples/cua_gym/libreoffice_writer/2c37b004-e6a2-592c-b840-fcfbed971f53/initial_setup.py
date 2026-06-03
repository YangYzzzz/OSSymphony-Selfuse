"""
Initial Setup: Create a Writer document with meeting notes.
Task ID: writer_fs_053
Domain: libreoffice_writer

The document contains meeting notes with the last two paragraphs holding
confidential information. No sections or protection exist yet.
The file is created as .docx via python-docx, then converted to .odt via
LibreOffice so that the agent works in ODF format (which preserves
Writer sections and protection natively).
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_053'
DOCX_TEMP = f'{WORKDIR}/{TASK_ID}_temp.docx'
OUTPUT = f'{WORKDIR}/{TASK_ID}.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    doc.add_heading('Quarterly Product Strategy Meeting', level=1)

    # --- Meeting metadata ---
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run('Date: ')
    run.bold = True
    run.font.size = Pt(11)
    meta.add_run('March 18, 2025').font.size = Pt(11)

    meta2 = doc.add_paragraph()
    meta2.paragraph_format.space_after = Pt(2)
    run = meta2.add_run('Location: ')
    run.bold = True
    run.font.size = Pt(11)
    meta2.add_run('Conference Room B, 3rd Floor').font.size = Pt(11)

    meta3 = doc.add_paragraph()
    meta3.paragraph_format.space_after = Pt(6)
    run = meta3.add_run('Attendees: ')
    run.bold = True
    run.font.size = Pt(11)
    meta3.add_run(
        'Sarah Chen (VP Product), Marcus Johnson (Engineering Lead), '
        'Priya Sharma (UX Director), David Kim (Marketing Manager), '
        'Elena Vasquez (QA Manager)'
    ).font.size = Pt(11)

    # --- Opening Remarks ---
    doc.add_heading('Opening Remarks', level=2)
    p = doc.add_paragraph(
        'Sarah Chen opened the meeting at 10:00 AM by reviewing the progress '
        'on the Q1 roadmap. She highlighted that 78% of planned features were '
        'delivered on schedule, with the remaining items deferred to Q2 due to '
        'resource constraints in the mobile development team.'
    )
    p.paragraph_format.space_after = Pt(6)

    # --- Engineering Update ---
    doc.add_heading('Engineering Update', level=2)
    p = doc.add_paragraph(
        'Marcus Johnson presented the engineering status report. The backend '
        'migration to microservices architecture is 65% complete. Performance '
        'benchmarks show a 40% improvement in API response times for the '
        'payment processing module. The team resolved 142 bugs in the last '
        'sprint cycle, bringing the critical bug count down to 7.'
    )
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph(
        'Marcus raised concerns about the upcoming database migration scheduled '
        'for April 5th. The migration window requires 4 hours of downtime, and '
        'he recommended scheduling it for Saturday between 2:00 AM and 6:00 AM '
        'to minimize customer impact. The team agreed unanimously.'
    )
    p2.paragraph_format.space_after = Pt(6)

    # --- UX Research Findings ---
    doc.add_heading('UX Research Findings', level=2)
    p = doc.add_paragraph(
        'Priya Sharma shared results from the latest usability study conducted '
        'with 35 participants across three user segments. Key findings include: '
        'the new onboarding flow reduced time-to-first-action by 28%, the '
        'redesigned dashboard received an average satisfaction score of 4.2 out '
        'of 5, and users requested a dark mode option (mentioned by 71% of '
        'participants).'
    )
    p.paragraph_format.space_after = Pt(6)

    # --- Marketing Pipeline ---
    doc.add_heading('Marketing Pipeline', level=2)
    p = doc.add_paragraph(
        'David Kim reported that the Q1 campaign generated 2,340 qualified leads, '
        'exceeding the target of 2,000 by 17%. The cost per acquisition dropped '
        'to $42.50 from $58.00 in the previous quarter. He proposed allocating '
        'an additional $15,000 to the social media budget for Q2 based on the '
        'strong ROI observed in Instagram and LinkedIn campaigns.'
    )
    p.paragraph_format.space_after = Pt(6)

    # --- Action Items ---
    doc.add_heading('Action Items', level=2)
    items = [
        'Marcus to finalize database migration plan and distribute to stakeholders by March 25',
        'Priya to deliver dark mode prototype mockups by April 1',
        'David to prepare revised Q2 marketing budget proposal for executive review',
        'Elena to schedule regression testing sprint for the new payment module',
        'Sarah to follow up with HR regarding the two open engineering positions',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Add spacing after bullets
    doc.add_paragraph('')

    # --- CONFIDENTIAL paragraphs (last two) ---
    p_conf1 = doc.add_paragraph(
        'The board has approved a preliminary budget of $2.4 million for the '
        'acquisition of DataSync Technologies. Due diligence is expected to '
        'conclude by April 30, 2025. This information must remain strictly '
        'confidential until the public announcement scheduled for May 15.'
    )
    p_conf1.paragraph_format.space_after = Pt(6)

    p_conf2 = doc.add_paragraph(
        'Additionally, the executive team has decided to restructure the '
        'Customer Success division, which will result in the elimination of '
        '12 positions and the creation of 8 new senior roles. Affected '
        'employees will be notified individually starting April 7. Any premature '
        'disclosure could have serious legal and reputational consequences.'
    )
    p_conf2.paragraph_format.space_after = Pt(6)

    doc.save(DOCX_TEMP)
    print(f'Temp docx created: {DOCX_TEMP}')

    # Convert to ODT using LibreOffice command-line
    subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'odt', '--outdir', WORKDIR, DOCX_TEMP],
        env={**os.environ, 'DISPLAY': ':0', 'HOME': WORKDIR},
        stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30
    )
    # The converted file will be writer_fs_053_temp.odt, rename it
    converted = f'{WORKDIR}/{TASK_ID}_temp.odt'
    if os.path.exists(converted):
        os.rename(converted, OUTPUT)
        os.remove(DOCX_TEMP)
        print(f'Initial file created (ODF): {OUTPUT}')
    else:
        # Fallback: keep docx
        os.rename(DOCX_TEMP, OUTPUT.replace('.odt', '.docx'))
        print(f'WARNING: ODF conversion failed, using docx')

    # GUI-ready: open the file in Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
