"""
Initial Setup: Mail merge letter template with contacts spreadsheet
Task ID: writer_mt_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_003'
OUTPUT_DOCX = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT_ODS = f'{DESKTOP}/contacts.ods'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_letter_template():
    """Create the thank-you letter with [ADDRESS HERE] placeholder."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Address placeholder at the top
    addr = doc.add_paragraph('[ADDRESS HERE]')
    addr.paragraph_format.space_after = Pt(24)
    run = addr.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Date line
    date_para = doc.add_paragraph('March 28, 2026')
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Salutation
    sal = doc.add_paragraph('Dear Valued Customer,')
    sal.paragraph_format.space_after = Pt(12)
    run = sal.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Body paragraph 1
    body1_text = (
        'Thank you for your recent purchase from Greenfield Home & Garden. '
        'We truly appreciate your business and hope that you are enjoying your new items. '
        'Your satisfaction is our top priority, and we want to make sure everything '
        'met your expectations.'
    )
    body1 = doc.add_paragraph(body1_text)
    body1.paragraph_format.space_after = Pt(12)
    run = body1.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Body paragraph 2
    body2_text = (
        'If you have any questions about your order or need assistance with any of our '
        'products, please do not hesitate to reach out. Our customer service team is '
        'available Monday through Friday from 8:00 AM to 6:00 PM EST at (555) 234-8900 '
        'or via email at support@greenfieldhg.com.'
    )
    body2 = doc.add_paragraph(body2_text)
    body2.paragraph_format.space_after = Pt(12)
    run = body2.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Body paragraph 3
    body3_text = (
        'As a token of our gratitude, we have enclosed a 15% discount coupon for your '
        'next purchase. We look forward to serving you again soon and hope to continue '
        'building a lasting relationship with you.'
    )
    body3 = doc.add_paragraph(body3_text)
    body3.paragraph_format.space_after = Pt(24)
    run = body3.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Closing
    closing = doc.add_paragraph('Warm regards,')
    closing.paragraph_format.space_after = Pt(36)
    run = closing.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    # Signature
    sig = doc.add_paragraph('Patricia Morales')
    sig.paragraph_format.space_after = Pt(0)
    run = sig.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)
    run.bold = True

    title_para = doc.add_paragraph('Customer Relations Manager')
    title_para.paragraph_format.space_after = Pt(0)
    run = title_para.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    company = doc.add_paragraph('Greenfield Home & Garden')
    run = company.runs[0]
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    doc.save(OUTPUT_DOCX)
    print(f'Initial letter template created: {OUTPUT_DOCX}')


def create_contacts_spreadsheet():
    """Create contacts.ods on the Desktop with 15 records."""
    import subprocess as sp

    # First create as xlsx, then convert to ods using LibreOffice
    # We use openpyxl to create an xlsx, then convert
    from openpyxl import Workbook

    os.makedirs(DESKTOP, exist_ok=True)
    temp_xlsx = f'/tmp/contacts.xlsx'

    wb = Workbook()
    ws = wb.active
    ws.title = 'Contacts'

    headers = ['Name', 'Street', 'City', 'State', 'ZipCode', 'Phone']
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)

    contacts = [
        ['Sarah Chen', '742 Maple Avenue', 'Portland', 'OR', '97201', '(503) 555-0142'],
        ['Marcus Johnson', '1285 Elm Street', 'Austin', 'TX', '78701', '(512) 555-0198'],
        ['Elena Rodriguez', '305 Oak Lane', 'Denver', 'CO', '80202', '(720) 555-0167'],
        ['David Kim', '89 Pine Road', 'Seattle', 'WA', '98101', '(206) 555-0234'],
        ['Amanda Foster', '1560 Birch Drive', 'Nashville', 'TN', '37201', '(615) 555-0189'],
        ['James Okafor', '423 Cedar Court', 'Chicago', 'IL', '60601', '(312) 555-0276'],
        ['Rachel Patel', '1792 Willow Way', 'San Diego', 'CA', '92101', '(619) 555-0153'],
        ['Thomas Wright', '651 Spruce Lane', 'Boston', 'MA', '02101', '(617) 555-0312'],
        ['Lisa Nakamura', '2104 Ash Street', 'Phoenix', 'AZ', '85001', '(480) 555-0245'],
        ['Robert Garcia', '378 Poplar Avenue', 'Miami', 'FL', '33101', '(305) 555-0198'],
        ['Jennifer Walsh', '925 Hickory Boulevard', 'Minneapolis', 'MN', '55401', '(612) 555-0167'],
        ['Michael Torres', '1437 Chestnut Drive', 'Atlanta', 'GA', '30301', '(404) 555-0289'],
        ['Samantha Lee', '562 Magnolia Court', 'Raleigh', 'NC', '27601', '(919) 555-0134'],
        ['Kevin Brown', '2890 Walnut Street', 'Philadelphia', 'PA', '19101', '(215) 555-0276'],
        ['Diana Petrov', '714 Sycamore Lane', 'Columbus', 'OH', '43201', '(614) 555-0198'],
    ]

    for r, row_data in enumerate(contacts, 2):
        for c, val in enumerate(row_data, 1):
            ws.cell(row=r, column=c, value=val)

    # Auto-fit column widths (approximate)
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 10
    ws.column_dimensions['F'].width = 16

    wb.save(temp_xlsx)
    print(f'Temporary xlsx created: {temp_xlsx}')

    # Convert to ODS using LibreOffice command line
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    result = sp.run(
        ['libreoffice', '--headless', '--convert-to', 'ods', '--outdir', DESKTOP, temp_xlsx],
        capture_output=True, text=True, env=env, timeout=30
    )
    print(f'LibreOffice convert stdout: {result.stdout}')
    print(f'LibreOffice convert stderr: {result.stderr}')

    if os.path.exists(OUTPUT_ODS):
        print(f'Contacts spreadsheet created: {OUTPUT_ODS}')
    else:
        # Fallback: just save as xlsx and rename (LibreOffice can open xlsx as well)
        import shutil
        fallback_path = f'{DESKTOP}/contacts.ods'
        # Try saving directly as ods using openpyxl is not possible, so we save xlsx
        # and LibreOffice on the VM should handle it
        shutil.copy(temp_xlsx, f'{DESKTOP}/contacts.xlsx')
        print(f'WARNING: ODS conversion failed, saved as xlsx at {DESKTOP}/contacts.xlsx')


def main():
    create_letter_template()
    create_contacts_spreadsheet()

    # Launch LibreOffice Writer with the letter template
    launch_gui(f'libreoffice --writer "{OUTPUT_DOCX}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
