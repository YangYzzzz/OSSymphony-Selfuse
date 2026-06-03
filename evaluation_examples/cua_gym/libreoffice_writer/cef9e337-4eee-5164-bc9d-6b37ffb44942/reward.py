"""
Reward Script: Insert a 5-column table with shaded header row and borders on all cells (8 data rows)
Task ID: osworld_writer_table_creation_003
Domain: libreoffice_writer
Scoring:
  - Component 1: Table exists with correct dimensions (5 cols x 9 rows)  — 0.5 pts
  - Component 2: Header row (row 0) has light gray background shading      — 0.3 pts
  - Component 3: Table uses 'Table Grid' style (borders on all cells)      — 0.2 pts
  Total: 1.0
"""

import os
from math import sqrt

# python-docx for .docx verification
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_table_creation_003'

# Light gray color tolerance: accept fills close to gray (R≈G≈B and not white/black)
def is_light_gray(fill_hex):
    """
    Check if a hex fill color is a light gray.
    Accepts colors where R, G, B channels are all >= 160 and close to each other.
    D3D3D3 (211,211,211) is the canonical light gray.
    """
    if not fill_hex or fill_hex.upper() in ('', 'NONE', 'AUTO', 'FFFFFF', 'FFFF00'):
        return False
    try:
        r = int(fill_hex[0:2], 16)
        g = int(fill_hex[2:4], 16)
        b = int(fill_hex[4:6], 16)
        # Check: all channels high (light), and close to each other (gray), not white (255,255,255)
        channel_avg = (r + g + b) / 3
        max_diff = max(abs(r - g), abs(g - b), abs(r - b))
        return channel_avg >= 150 and max_diff <= 40 and not (r >= 250 and g >= 250 and b >= 250)
    except Exception:
        return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Task: Insert a 5-column table with shaded header row and borders on all cells (8 data rows).
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: Table exists with correct dimensions (0.5 points) ---
    # A 5-column x 9-row table (1 header + 8 data rows) must be present.
    # This FAILS on initial (no tables), PASSES on golden.
    try:
        tables = doc.tables
        if len(tables) == 0:
            print("FAIL: Component 1 — No tables found in document")
        else:
            # Find the qualifying table
            qualifying_table = None
            for t in tables:
                num_rows = len(t.rows)
                num_cols = len(t.columns)
                if num_cols == 5 and num_rows == 9:
                    qualifying_table = t
                    break

            if qualifying_table is not None:
                print(f"PASS: Component 1 — Found table with 5 columns x 9 rows (0.5 pts)")
                total_score += 0.5
            else:
                # Partial: table exists but dimensions wrong
                best_table = tables[0]
                num_rows = len(best_table.rows)
                num_cols = len(best_table.columns)
                print(f"FAIL: Component 1 — Expected 5 cols x 9 rows, found {num_cols} cols x {num_rows} rows")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: Header row has light gray background shading (0.3 points) ---
    # All 5 cells in row 0 should have a light gray fill (e.g., D3D3D3).
    # This FAILS on initial (no table), PASSES on golden.
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No table to check header shading")
        else:
            table = doc.tables[0]
            # Use the qualifying table if found, otherwise first table
            for t in doc.tables:
                if len(t.columns) == 5 and len(t.rows) >= 1:
                    table = t
                    break

            header_row = table.rows[0]
            shaded_cells = 0
            total_cells = len(header_row.cells)
            for cell in header_row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill_val = shd.get(qn('w:fill'))
                        if fill_val and is_light_gray(fill_val):
                            shaded_cells += 1

            if shaded_cells == total_cells and total_cells == 5:
                print(f"PASS: Component 2 — All 5 header cells have light gray shading (0.3 pts)")
                total_score += 0.3
            elif shaded_cells > 0:
                print(f"FAIL: Component 2 — Only {shaded_cells}/{total_cells} header cells have light gray shading")
            else:
                print(f"FAIL: Component 2 — No header cells have light gray shading (shaded_cells=0)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: Table uses 'Table Grid' style (borders on all cells) (0.2 points) ---
    # 'Table Grid' is the standard style that renders borders on all cells.
    # This FAILS on initial (no table), PASSES on golden.
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No table to check border style")
        else:
            table = doc.tables[0]
            for t in doc.tables:
                if len(t.columns) == 5 and len(t.rows) >= 1:
                    table = t
                    break

            style_name = table.style.name if table.style else ''
            # Accept 'Table Grid' or any style with 'grid' in the name (case-insensitive)
            has_grid_style = 'grid' in style_name.lower() or 'Table Grid' in style_name
            if has_grid_style:
                print(f"PASS: Component 3 — Table uses border style '{style_name}' (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Expected 'Table Grid' style, found '{style_name}'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
