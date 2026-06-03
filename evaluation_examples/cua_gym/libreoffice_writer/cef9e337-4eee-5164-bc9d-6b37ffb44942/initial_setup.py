"""
Initial Setup: Project tracking report document (no tables)
Task ID: osworld_writer_table_creation_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_table_creation_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Q2 2025 Project Tracking Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introduction ---
    doc.add_heading('Executive Summary', level=1)
    intro = doc.add_paragraph(
        'This report provides an overview of all active projects for Q2 2025. '
        'It summarizes current status, key milestones, resource allocations, and '
        'outstanding blockers across the engineering and product teams. '
        'Project leads are expected to update their respective task status by the end of each sprint.'
    )

    doc.add_paragraph(
        'The portfolio consists of eight active work streams spanning infrastructure upgrades, '
        'new feature development, customer-facing integrations, and compliance initiatives. '
        'All items below have been reviewed and approved by the Program Management Office (PMO).'
    )

    # --- Project Overview section ---
    doc.add_heading('Project Overview', level=1)
    doc.add_paragraph(
        'The following section outlines each project\'s current status and key tracking metrics. '
        'Priorities have been assigned based on business impact, regulatory deadlines, and '
        'available engineering capacity. Please refer to the task tracker for detailed subtask breakdowns.'
    )

    # --- Team Leads section ---
    doc.add_heading('Team Leads and Responsibilities', level=1)

    leads_intro = doc.add_paragraph(
        'Each project is assigned a dedicated lead responsible for coordination, '
        'status reporting, and escalation management:'
    )

    leads = [
        'Infrastructure Modernization — Lead: Priya Nair (Backend Platform)',
        'Customer Portal Redesign — Lead: James Whitfield (Frontend)',
        'Data Pipeline Migration — Lead: Elena Sorokina (Data Engineering)',
        'Security Compliance Audit — Lead: Carlos Mendez (InfoSec)',
        'Mobile SDK Release — Lead: Aiko Tanaka (Mobile)',
        'Analytics Dashboard v2 — Lead: Daniel Osei (Product Analytics)',
        'API Gateway Upgrade — Lead: Fatima Al-Rashid (Platform)',
        'Documentation Overhaul — Lead: Liam Brennan (Technical Writing)',
    ]

    for lead in leads:
        doc.add_paragraph(lead, style='List Bullet')

    # --- Status Definitions section ---
    doc.add_heading('Status Definitions', level=1)

    doc.add_paragraph(
        'The following status labels are used throughout this document:'
    )

    statuses = [
        'On Track — Project is proceeding according to schedule with no critical blockers.',
        'At Risk — One or more dependencies are delayed; mitigation plan in progress.',
        'Blocked — Project cannot proceed without resolution of a critical dependency.',
        'Completed — All deliverables have been submitted and approved.',
        'In Review — Work is complete pending final sign-off from stakeholders.',
    ]

    for s in statuses:
        doc.add_paragraph(s, style='List Bullet')

    # --- Task Breakdown section (NO TABLE - that is the task) ---
    doc.add_heading('Task Breakdown', level=1)

    doc.add_paragraph(
        'The detailed task breakdown table will be inserted below. '
        'Each row represents a discrete work item with assigned owner, current status, '
        'due date, and priority level. Please insert the project task tracking table '
        'with appropriate column headers and formatting as outlined in the reporting guidelines.'
    )

    # --- Notes section ---
    doc.add_heading('Notes and Escalations', level=1)

    doc.add_paragraph(
        'Any items marked as Blocked or At Risk must be escalated to the PMO within 24 hours. '
        'Weekly sync meetings are held every Tuesday at 10:00 AM PST. '
        'Please ensure all updates are reflected in this document before the Monday EOD deadline.'
    )

    doc.add_paragraph(
        'For questions regarding resource allocation or timeline adjustments, '
        'contact the PMO at pmo@company.internal or reach out directly to your assigned program manager.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
