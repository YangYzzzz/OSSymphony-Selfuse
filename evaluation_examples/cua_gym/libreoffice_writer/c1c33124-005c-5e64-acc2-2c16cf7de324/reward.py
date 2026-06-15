"""
Reward Script: Insert column break before 'Furthermore' paragraph
Task ID: wrpara_038
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): Column break exists in/before the 'Furthermore' paragraph
  Component 2 (0.4): Exactly one column break, 3-column layout preserved, text intact (compound)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'wrpara_038'
WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WML_NS}


def find_column_breaks(doc):
    """Find all column breaks in the document. Returns list of (para_index, run_index) tuples."""
    breaks = []
    for pidx, para in enumerate(doc.paragraphs):
        for ridx, run in enumerate(para.runs):
            for br in run.element.findall('.//w:br', NS):
                br_type = br.attrib.get(f'{{{WML_NS}}}type', 'line')
                if br_type == 'column':
                    breaks.append((pidx, ridx))
    return breaks


def find_furthermore_para(doc):
    """Find the paragraph index that starts with 'Furthermore'."""
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('Furthermore'):
            return idx
    return None


def get_column_count(doc):
    """Get the number of columns from the first section."""
    try:
        section = doc.sections[0]
        cols_elems = section._sectPr.findall('.//w:cols', NS)
        if cols_elems:
            num = cols_elems[0].attrib.get(f'{{{WML_NS}}}num', '1')
            return int(num)
    except Exception:
        pass
    return 1


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the 'Furthermore' paragraph
    furthermore_idx = find_furthermore_para(doc)
    if furthermore_idx is None:
        print("CRITICAL: No paragraph starting with 'Furthermore' found in document")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: 'Furthermore' paragraph found at index {furthermore_idx}")

    # Find all column breaks
    col_breaks = find_column_breaks(doc)
    print(f"INFO: Found {len(col_breaks)} column break(s): {col_breaks}")

    # Component 1: Column break exists in/before the 'Furthermore' paragraph (0.6 points)
    # The column break should be within that paragraph (before the text), or at the
    # end of the preceding paragraph. This is the core task requirement.
    comp1_passed = len([b for b in col_breaks if b[0] == furthermore_idx]) > 0
    comp1_alt = len([b for b in col_breaks if b[0] == furthermore_idx - 1]) > 0
    try:
        if comp1_passed:
            print(f"PASS: Component 1 -- Column break found IN 'Furthermore' paragraph (idx {furthermore_idx}) (0.6 pts)")
            total_score += 0.6
        elif comp1_alt:
            print(f"PASS: Component 1 -- Column break found at end of paragraph before 'Furthermore' (idx {furthermore_idx - 1}) (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 -- No column break at or just before 'Furthermore' paragraph. Breaks found at: {col_breaks}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Exactly one column break AND 3-column layout AND text intact (0.4 points)
    # This is a compound check anchored to the column break: only awards points if
    # the column break is correctly placed AND additional quality criteria are met.
    # This ensures initial_env (no column break) scores 0.
    try:
        if not (comp1_passed or comp1_alt):
            print(f"FAIL: Component 2 -- Skipped because no correct column break was found (prerequisite)")
        else:
            col_count = get_column_count(doc)
            furthermore_text = doc.paragraphs[furthermore_idx].text.strip()

            sub_checks_passed = 0
            sub_total = 2

            # Sub-check 2a: exactly one column break in the document
            if len(col_breaks) == 1:
                sub_checks_passed += 1
                print(f"  Sub-check 2a PASS: Exactly 1 column break in document")
            else:
                print(f"  Sub-check 2a FAIL: Expected 1 column break, found {len(col_breaks)}")

            # Sub-check 2b: 3-column layout preserved and text intact
            cols_ok = (col_count == 3)
            text_ok = (furthermore_text.startswith('Furthermore') and len(furthermore_text) > 50)
            if cols_ok and text_ok:
                sub_checks_passed += 1
                print(f"  Sub-check 2b PASS: 3-column layout (cols={col_count}), text intact ({len(furthermore_text)} chars)")
            else:
                if not cols_ok:
                    print(f"  Sub-check 2b FAIL: Column count is {col_count}, expected 3")
                if not text_ok:
                    print(f"  Sub-check 2b FAIL: 'Furthermore' text issue: {furthermore_text[:30]}..., len={len(furthermore_text)}")

            comp2_score = round(0.4 * (sub_checks_passed / sub_total), 2)
            if comp2_score > 0:
                print(f"PASS: Component 2 -- {sub_checks_passed}/{sub_total} sub-checks passed ({comp2_score} pts)")
                total_score += comp2_score
            else:
                print(f"FAIL: Component 2 -- 0/{sub_total} sub-checks passed")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    final_score = round(min(total_score, 1.0), 1)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
