"""
Initial Setup: Insert a column break in a 3-column brochure document
Task ID: wrpara_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'wrpara_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page with narrow margins for 3-column layout
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Configure 3 columns via XML
    sectPr = section._sectPr
    cols = sectPr.makeelement(qn('w:cols'), {
        qn('w:num'): '3',
        qn('w:space'): '360',  # 0.25 inch spacing between columns
        qn('w:equalWidth'): '1',
    })
    # Remove any existing cols element
    for existing in sectPr.findall(qn('w:cols')):
        sectPr.remove(existing)
    sectPr.append(cols)

    # --- Brochure Title ---
    title = doc.add_heading('2025 Annual Market Research Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by Meridian Analytics Group')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Body paragraphs: enough text to fill columns naturally ---
    # Column 1 content
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    run = p1.add_run('Executive Summary')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'The global consumer electronics market experienced a significant '
        'transformation during the fiscal year 2024-2025. Driven by rapid '
        'advancements in artificial intelligence and shifting consumer '
        'preferences, the sector recorded unprecedented growth across multiple '
        'product categories. Our comprehensive analysis spans 47 countries '
        'and covers data from over 12,000 retail outlets.',
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'Revenue across all segments reached $487.3 billion, representing '
        'a 14.2% increase compared to the prior year. The Asia-Pacific region '
        'continued to dominate, accounting for 41% of total sales, while '
        'North America and Europe contributed 28% and 22% respectively. '
        'Emerging markets in Latin America and Sub-Saharan Africa showed '
        'the fastest growth rates at 23.7% and 19.4%.',
    ).paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run = p2.add_run('Key Market Drivers')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Several macroeconomic factors contributed to this growth trajectory. '
        'The widespread adoption of 5G networks enabled new product categories '
        'including connected home devices and wearable health monitors. '
        'Additionally, the post-pandemic shift toward hybrid work environments '
        'sustained demand for productivity hardware such as laptops, monitors, '
        'and ergonomic peripherals.',
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'Consumer spending on smart home technology surged by 31% year-over-year, '
        'with voice-activated assistants and automated security systems leading '
        'the category. The average household now owns 8.4 connected devices, up '
        'from 6.1 in the previous year. Premium smartphone sales grew by 18%, '
        'driven largely by camera improvements and AI-powered features that '
        'appealed to content creators and professional users alike.',
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'The enterprise segment proved equally robust. Cloud computing infrastructure '
        'investments rose by 27%, and semiconductor demand outpaced supply for the '
        'third consecutive quarter. Companies across all industries increased their '
        'technology budgets by an average of 16%, reflecting the ongoing digital '
        'transformation imperative.',
    ).paragraph_format.space_after = Pt(6)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    run = p3.add_run('Regional Performance')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'In the Asia-Pacific region, China maintained its position as the largest '
        'single market, with consumer electronics revenue reaching $128.6 billion. '
        'India emerged as the fastest-growing major market, recording 29.3% growth '
        'driven by expanding middle-class purchasing power and improved logistics '
        'infrastructure. Japan and South Korea continued to lead in innovation, '
        'particularly in display technology and robotics.',
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'North American markets benefited from strong holiday season sales and '
        'the successful launch of several flagship products. The United States '
        'accounted for $112.4 billion in revenue, while Canada contributed '
        '$24.8 billion. Both markets showed particular strength in subscription-based '
        'services and digital content platforms, which grew by 35% and now represent '
        'a significant recurring revenue stream for manufacturers.',
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'European markets demonstrated resilience despite ongoing economic '
        'headwinds. Germany, France, and the United Kingdom collectively generated '
        '$78.2 billion in revenue. The region showed particular appetite for '
        'sustainable and energy-efficient products, with eco-certified electronics '
        'outselling conventional alternatives by a 2:1 margin in Scandinavian '
        'countries.',
    ).paragraph_format.space_after = Pt(6)

    # This is the critical paragraph that must appear mid-column-2
    # and will be moved to column 3 top by the column break task
    doc.add_paragraph(
        'Furthermore, our analysis shows that consumer loyalty patterns have '
        'shifted dramatically over the past twelve months. Brand switching rates '
        'increased by 22% across all demographics, with younger consumers aged '
        '18-34 showing the highest propensity to explore alternative brands. '
        'Price sensitivity has intensified, and 67% of surveyed consumers now '
        'compare at least four different products before making a purchase decision.',
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'The sustainability factor has become a decisive purchasing criterion '
        'for 43% of consumers globally, up from 28% two years ago. Manufacturers '
        'who invested in recyclable packaging and carbon-neutral production '
        'processes reported 15% higher customer retention rates. This trend is '
        'expected to accelerate as regulatory frameworks around electronic waste '
        'become more stringent in major markets.',
    ).paragraph_format.space_after = Pt(6)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    run = p4.add_run('Outlook and Recommendations')
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Looking ahead to 2026, we project continued growth at a rate of 11-13% '
        'driven by emerging technology categories including augmented reality '
        'glasses, advanced wearable health devices, and AI-powered personal '
        'assistants. Supply chain stabilization and expanding manufacturing '
        'capacity should alleviate component shortages that constrained growth '
        'in previous periods.',
    ).paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'We recommend that industry stakeholders prioritize investment in three '
        'key areas: artificial intelligence integration, sustainability compliance, '
        'and direct-to-consumer distribution channels. Companies that successfully '
        'address all three dimensions will be best positioned to capture market '
        'share in an increasingly competitive landscape.',
    ).paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
