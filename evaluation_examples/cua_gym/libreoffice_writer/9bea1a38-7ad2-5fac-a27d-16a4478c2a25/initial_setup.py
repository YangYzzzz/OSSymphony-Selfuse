"""
Initial Setup: Create a thesis document for PDF/A export task
Task ID: writer_acad_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title Page ---
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading('', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Machine Learning Approaches for Early Detection\nof Neurodegenerative Diseases:\nA Multi-Modal Imaging Study')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy in Computational Neuroscience')
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Elena Vasquez-Rodriguez')
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run('Department of Biomedical Engineering\nStanford University\nMarch 2025')
    run.font.size = Pt(12)

    doc.add_paragraph()

    committee = doc.add_paragraph()
    committee.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = committee.add_run('Thesis Committee:\nProf. James K. Mitchell (Chair)\nProf. Ananya Sharma\nProf. Robert T. Nakamura\nDr. Lisa Fernandez-Ortiz')
    run.font.size = Pt(11)

    # --- Page Break: Abstract ---
    doc.add_page_break()

    doc.add_heading('Abstract', level=1)
    abstract_text = (
        'This dissertation presents a novel computational framework for the early detection '
        'of neurodegenerative diseases, specifically Alzheimer\'s disease (AD) and Parkinson\'s '
        'disease (PD), through the integration of multi-modal neuroimaging data. Our approach '
        'combines structural magnetic resonance imaging (sMRI), diffusion tensor imaging (DTI), '
        'and functional MRI (fMRI) with advanced deep learning architectures to identify '
        'disease-specific biomarkers during the prodromal phase, up to five years before '
        'clinical symptom onset.'
    )
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.space_after = Pt(12)

    abstract_text2 = (
        'We developed a multi-stream convolutional neural network (MS-CNN) that processes '
        'each imaging modality through dedicated encoder pathways before fusing representations '
        'in a shared latent space. The model was trained and validated on the Alzheimer\'s Disease '
        'Neuroimaging Initiative (ADNI) dataset comprising 2,847 participants and further tested '
        'on the Parkinson\'s Progression Markers Initiative (PPMI) cohort of 1,234 subjects. '
        'Our framework achieved a classification accuracy of 94.3% for AD detection (sensitivity: '
        '92.1%, specificity: 96.5%) and 89.7% for PD detection (sensitivity: 87.4%, specificity: '
        '91.8%), significantly outperforming single-modality approaches and existing state-of-the-art '
        'methods.'
    )
    p = doc.add_paragraph(abstract_text2)
    p.paragraph_format.space_after = Pt(12)

    keywords = doc.add_paragraph()
    run = keywords.add_run('Keywords: ')
    run.bold = True
    keywords.add_run(
        'neurodegenerative diseases, deep learning, multi-modal neuroimaging, '
        'Alzheimer\'s disease, Parkinson\'s disease, early detection, biomarkers'
    )

    # --- Page Break: Table of Contents placeholder ---
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('Abstract', 'ii'),
        ('List of Figures', 'v'),
        ('List of Tables', 'vi'),
        ('Chapter 1: Introduction', '1'),
        ('  1.1 Background and Motivation', '2'),
        ('  1.2 Research Objectives', '5'),
        ('  1.3 Thesis Organization', '7'),
        ('Chapter 2: Literature Review', '9'),
        ('  2.1 Neurodegenerative Disease Pathology', '9'),
        ('  2.2 Neuroimaging Modalities', '15'),
        ('  2.3 Machine Learning in Medical Imaging', '22'),
        ('Chapter 3: Methodology', '31'),
        ('  3.1 Data Acquisition and Preprocessing', '31'),
        ('  3.2 Multi-Stream CNN Architecture', '38'),
        ('  3.3 Training Protocol', '45'),
        ('Chapter 4: Results', '52'),
        ('  4.1 Classification Performance', '52'),
        ('  4.2 Ablation Studies', '60'),
        ('  4.3 Biomarker Analysis', '67'),
        ('Chapter 5: Discussion', '74'),
        ('Chapter 6: Conclusion', '82'),
        ('References', '86'),
        ('Appendix A: Supplementary Figures', '95'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run(f'\t{page}')

    # --- Page Break: Chapter 1 ---
    doc.add_page_break()
    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('1.1 Background and Motivation', level=2)
    doc.add_paragraph(
        'Neurodegenerative diseases represent one of the most significant public health '
        'challenges of the twenty-first century. According to the World Health Organization, '
        'approximately 55 million people worldwide currently live with dementia, with nearly '
        '10 million new cases diagnosed annually (WHO, 2023). Alzheimer\'s disease alone accounts '
        'for 60-70% of all dementia cases, while Parkinson\'s disease affects over 8.5 million '
        'individuals globally. The economic burden is staggering: the global cost of dementia care '
        'exceeded $1.3 trillion in 2024, and projections indicate this figure will surpass $2.8 '
        'trillion by 2030 (Wimo et al., 2024).'
    )

    doc.add_paragraph(
        'A critical challenge in managing neurodegenerative diseases lies in the temporal gap '
        'between pathological onset and clinical diagnosis. In Alzheimer\'s disease, amyloid-beta '
        'plaque accumulation and tau protein hyperphosphorylation begin 15-20 years before cognitive '
        'symptoms manifest (Jack et al., 2013). Similarly, Parkinson\'s disease involves progressive '
        'dopaminergic neuron loss in the substantia nigra that commences years before motor symptoms '
        'become clinically apparent (Postuma et al., 2012). This prolonged preclinical phase offers '
        'a therapeutic window during which disease-modifying interventions could be most effective, '
        'provided early and accurate detection methods are available.'
    )

    doc.add_paragraph(
        'Recent advances in neuroimaging technology have provided unprecedented windows into '
        'brain structure and function. Structural MRI reveals volumetric changes in cortical and '
        'subcortical regions, diffusion tensor imaging maps white matter tract integrity, and '
        'functional MRI captures neural activity patterns through blood-oxygen-level-dependent '
        '(BOLD) signal changes. Each modality captures distinct but complementary aspects of '
        'neurodegeneration. However, the challenge lies in effectively integrating these '
        'heterogeneous data streams to extract clinically meaningful patterns that precede '
        'symptomatic presentation.'
    )

    doc.add_heading('1.2 Research Objectives', level=2)
    doc.add_paragraph(
        'This thesis addresses the fundamental challenge of early neurodegenerative disease '
        'detection through a unified computational framework. Our primary research objectives are:'
    )
    objectives = [
        'To develop a multi-stream deep learning architecture capable of processing and '
        'integrating heterogeneous neuroimaging modalities (sMRI, DTI, and fMRI) within a '
        'unified computational framework.',
        'To identify modality-specific and cross-modal biomarkers that reliably distinguish '
        'prodromal neurodegenerative disease states from healthy aging, with particular focus '
        'on sensitivity to early pathological changes.',
        'To evaluate the clinical utility of the proposed framework through rigorous validation '
        'on large-scale, multi-center datasets, assessing generalizability across demographic '
        'groups and imaging acquisition protocols.',
        'To investigate the interpretability of learned representations using gradient-based '
        'attribution methods, providing insights into the neuroanatomical substrates that drive '
        'model predictions and their alignment with known disease pathology.',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Number')

    doc.add_heading('1.3 Thesis Organization', level=2)
    doc.add_paragraph(
        'The remainder of this thesis is organized as follows. Chapter 2 provides a comprehensive '
        'review of the relevant literature, spanning neurodegenerative disease pathophysiology, '
        'neuroimaging methodologies, and machine learning approaches in medical imaging. Chapter 3 '
        'details our proposed methodology, including data acquisition pipelines, preprocessing '
        'steps, network architecture design, and training procedures. Chapter 4 presents '
        'experimental results, encompassing classification performance metrics, ablation studies, '
        'and biomarker analyses. Chapter 5 discusses the implications of our findings, limitations, '
        'and potential clinical applications. Finally, Chapter 6 summarizes our contributions and '
        'outlines directions for future research.'
    )

    # --- Page Break: Chapter 2 ---
    doc.add_page_break()
    doc.add_heading('Chapter 2: Literature Review', level=1)

    doc.add_heading('2.1 Neurodegenerative Disease Pathology', level=2)
    doc.add_paragraph(
        'The pathological hallmarks of Alzheimer\'s disease include extracellular amyloid-beta '
        '(A-beta) plaque deposits and intracellular neurofibrillary tangles composed of '
        'hyperphosphorylated tau protein (Braak & Braak, 1991). The amyloid cascade hypothesis '
        'posits that A-beta accumulation initiates a pathological cascade leading to synaptic '
        'dysfunction, neuroinflammation, and ultimately neuronal death (Hardy & Higgins, 1992). '
        'Recent evidence from longitudinal imaging studies suggests that A-beta deposition begins '
        'in the default mode network regions, particularly the precuneus and posterior cingulate '
        'cortex, before spreading to frontal and temporal areas following a characteristic spatial '
        'pattern (Palmqvist et al., 2017).'
    )

    doc.add_paragraph(
        'Parkinson\'s disease is characterized by the progressive loss of dopaminergic neurons '
        'in the substantia nigra pars compacta and the formation of Lewy bodies containing '
        'misfolded alpha-synuclein protein (Spillantini et al., 1997). Braak staging of PD '
        'pathology suggests a caudal-to-rostral progression, beginning in the dorsal motor '
        'nucleus of the vagus nerve and olfactory bulb before ascending through the brainstem '
        'to affect the substantia nigra and eventually the neocortex (Braak et al., 2003). '
        'Non-motor symptoms including anosmia, REM sleep behavior disorder, and constipation '
        'often precede motor manifestations by years, consistent with this staging model.'
    )

    doc.add_heading('2.2 Neuroimaging Modalities', level=2)
    doc.add_paragraph(
        'Structural MRI has been the workhorse of neuroimaging-based dementia research for '
        'decades. Voxel-based morphometry (VBM) studies have consistently demonstrated '
        'hippocampal atrophy as a sensitive marker of AD progression, with annual volume loss '
        'rates of 3-5% in AD patients compared to 0.5-1% in cognitively normal elderly '
        '(Frisoni et al., 2010). FreeSurfer-derived cortical thickness measurements have '
        'further revealed entorhinal cortex thinning as one of the earliest structural changes '
        'detectable in preclinical AD (Dickerson et al., 2009).'
    )

    doc.add_paragraph(
        'Diffusion tensor imaging provides unique sensitivity to white matter microstructural '
        'changes through metrics including fractional anisotropy (FA), mean diffusivity (MD), '
        'and tractography-based structural connectivity measures. In AD, decreased FA and '
        'increased MD have been reported in the cingulum bundle, fornix, and corpus callosum, '
        'reflecting Wallerian degeneration secondary to cortical neuron loss (Acosta-Cabronero '
        'et al., 2010). In PD, DTI has revealed alterations in nigrostriatal pathways and '
        'interhemispheric connections even in early disease stages (Vaillancourt et al., 2009).'
    )

    # --- Page Break: Chapter 3 ---
    doc.add_page_break()
    doc.add_heading('Chapter 3: Methodology', level=1)

    doc.add_heading('3.1 Data Acquisition and Preprocessing', level=2)
    doc.add_paragraph(
        'Data for this study were obtained from two large-scale multi-center research databases. '
        'The primary dataset was drawn from the Alzheimer\'s Disease Neuroimaging Initiative (ADNI; '
        'adni.loni.usc.edu), which provided longitudinal multi-modal imaging data for 2,847 '
        'participants across four diagnostic categories: cognitively normal (CN, n=814), subjective '
        'cognitive decline (SCD, n=326), mild cognitive impairment (MCI, n=1,142), and Alzheimer\'s '
        'disease (AD, n=565). The secondary validation dataset was obtained from the Parkinson\'s '
        'Progression Markers Initiative (PPMI; ppmi-info.org), comprising 1,234 participants '
        'including healthy controls (HC, n=423), SWEDD (n=178), and PD patients (n=633).'
    )

    # Data table
    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    headers = ['Cohort', 'N', 'Age (mean +/- SD)', 'Female %', 'MMSE (mean)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data_rows = [
        ['CN (ADNI)', '814', '73.2 +/- 6.8', '52.1%', '29.1'],
        ['SCD (ADNI)', '326', '71.8 +/- 7.1', '54.3%', '28.7'],
        ['MCI (ADNI)', '1,142', '74.1 +/- 7.4', '43.8%', '27.2'],
        ['AD (ADNI)', '565', '75.6 +/- 8.2', '48.7%', '21.4'],
        ['PD (PPMI)', '633', '62.4 +/- 9.7', '37.2%', '27.8'],
    ]
    for r, row_data in enumerate(data_rows, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    caption = doc.add_paragraph()
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption.add_run('Table 3.1: Demographic characteristics of study participants')
    run.italic = True
    run.font.size = Pt(10)

    doc.add_heading('3.2 Multi-Stream CNN Architecture', level=2)
    doc.add_paragraph(
        'Our proposed Multi-Stream Convolutional Neural Network (MS-CNN) architecture processes '
        'each imaging modality through a dedicated encoder stream before fusing the learned '
        'representations in a shared latent space. The sMRI stream employs a 3D ResNet-50 backbone '
        'pretrained on the UK Biobank dataset, processing skull-stripped, MNI-registered T1-weighted '
        'volumes at 1mm isotropic resolution. The DTI stream utilizes a custom 3D encoder '
        'operating on 6-channel diffusion tensor images (FA, MD, axial diffusivity, radial '
        'diffusivity, and two principal eigenvector components). The fMRI stream processes '
        'temporal correlation matrices derived from 116 ROIs defined by the AAL atlas through '
        'a graph attention network (GAT) architecture.'
    )

    doc.add_paragraph(
        'Feature fusion is accomplished through a cross-modal attention mechanism that learns '
        'inter-modality dependencies. Each stream produces a 512-dimensional feature vector, '
        'which is projected into a shared 256-dimensional space through learned linear '
        'transformations. Multi-head cross-attention (8 heads, key/query dimension 32) enables '
        'each modality to attend to relevant features from other modalities. The fused '
        'representation is passed through two fully-connected layers (256 -> 128 -> C, where C '
        'is the number of diagnostic categories) with dropout (p=0.3) and batch normalization.'
    )

    doc.add_heading('3.3 Training Protocol', level=2)
    doc.add_paragraph(
        'Models were trained using a stratified 5-fold cross-validation scheme with an 80/10/10 '
        'train/validation/test split, ensuring balanced representation of diagnostic categories '
        'within each fold. We employed the AdamW optimizer with an initial learning rate of '
        '1e-4 and cosine annealing schedule (T_max=50 epochs). Training was performed on 4 '
        'NVIDIA A100 GPUs with a per-GPU batch size of 8, yielding an effective batch size of '
        '32. Mixed-precision training (FP16) was used to reduce memory consumption and accelerate '
        'computation. Class imbalance was addressed through focal loss (gamma=2.0, alpha derived '
        'from inverse class frequencies) combined with oversampling of minority classes during '
        'training.'
    )

    # --- Page Break: Chapter 4 ---
    doc.add_page_break()
    doc.add_heading('Chapter 4: Results', level=1)

    doc.add_heading('4.1 Classification Performance', level=2)
    doc.add_paragraph(
        'Table 4.1 presents the classification performance of our MS-CNN framework compared '
        'to baseline single-modality models and existing state-of-the-art methods on the ADNI '
        'test set. Our multi-modal approach achieved an overall accuracy of 94.3% for the '
        'binary CN vs. AD classification task, with a sensitivity of 92.1% and specificity of '
        '96.5%. For the more challenging four-class classification (CN/SCD/MCI/AD), the model '
        'achieved a macro-averaged F1-score of 0.847, representing a 7.2% improvement over the '
        'best-performing single-modality baseline (sMRI ResNet-50, F1=0.775).'
    )

    # Results table
    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    headers2 = ['Method', 'Accuracy (%)', 'Sensitivity (%)', 'Specificity (%)', 'AUC']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    results_data = [
        ['sMRI Only', '87.2', '84.6', '89.8', '0.923'],
        ['DTI Only', '79.5', '76.3', '82.7', '0.867'],
        ['fMRI Only', '82.1', '79.8', '84.4', '0.891'],
        ['Liu et al. (2023)', '91.8', '89.4', '94.2', '0.961'],
        ['MS-CNN (Ours)', '94.3', '92.1', '96.5', '0.982'],
    ]
    for r, row_data in enumerate(results_data, 1):
        for c, val in enumerate(row_data):
            cell = table2.cell(r, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            if r == 5:  # Bold our results
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    caption2 = doc.add_paragraph()
    caption2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption2.add_run('Table 4.1: Binary classification performance (CN vs. AD) on ADNI test set')
    run.italic = True
    run.font.size = Pt(10)

    # --- Page Break: Chapter 5 ---
    doc.add_page_break()
    doc.add_heading('Chapter 5: Discussion', level=1)
    doc.add_paragraph(
        'The results presented in this thesis demonstrate that multi-modal neuroimaging integration '
        'through deep learning provides substantial improvements in early neurodegenerative disease '
        'detection compared to single-modality approaches. Our MS-CNN framework\'s superior '
        'performance can be attributed to three key factors: (1) the complementary nature of '
        'structural, microstructural, and functional imaging data; (2) the cross-modal attention '
        'mechanism that learns meaningful inter-modality relationships; and (3) the use of '
        'modality-specific preprocessing pipelines optimized for each data type.'
    )

    doc.add_paragraph(
        'The cross-modal attention maps revealed several clinically meaningful patterns. For AD '
        'classification, the model learned to jointly attend to hippocampal volume from sMRI, '
        'fornix FA from DTI, and default mode network connectivity from fMRI. This multi-modal '
        'signature aligns with the known pathological cascade in AD, where medial temporal lobe '
        'atrophy, limbic tract degeneration, and network disconnection occur in parallel during '
        'disease progression (Jack et al., 2013). For PD detection, the model prioritized '
        'substantia nigra morphology, nigrostriatal FA, and motor network functional connectivity, '
        'consistent with the dopaminergic degeneration that underlies motor symptomatology.'
    )

    doc.add_paragraph(
        'Several limitations of this work should be acknowledged. First, while ADNI and PPMI '
        'represent large, well-characterized cohorts, they may not fully capture the diversity '
        'of real-world clinical populations, particularly with respect to racial and ethnic '
        'representation, comorbidities, and imaging acquisition variability. Second, our current '
        'framework requires all three imaging modalities to be available, which may limit clinical '
        'applicability in resource-constrained settings. Future work should explore modality-agnostic '
        'architectures that can gracefully handle missing data. Third, the computational requirements '
        'of the full MS-CNN pipeline may present challenges for real-time clinical deployment, '
        'though model distillation and optimization techniques could address this concern.'
    )

    # --- Page Break: Chapter 6 ---
    doc.add_page_break()
    doc.add_heading('Chapter 6: Conclusion', level=1)
    doc.add_paragraph(
        'This thesis presented a comprehensive computational framework for early detection of '
        'neurodegenerative diseases through multi-modal neuroimaging integration. Our Multi-Stream '
        'CNN architecture with cross-modal attention fusion achieved state-of-the-art performance '
        'on both Alzheimer\'s disease and Parkinson\'s disease detection tasks, demonstrating the '
        'value of combining complementary imaging modalities within a unified deep learning '
        'framework. The interpretability analyses provided clinically meaningful insights into '
        'the neuroanatomical substrates driving model predictions, enhancing confidence in the '
        'biological plausibility of the learned representations.'
    )

    doc.add_paragraph(
        'Future research directions include: (1) extension to additional neurodegenerative '
        'conditions including frontotemporal dementia and Lewy body dementia; (2) integration '
        'of genetic risk factors (APOE genotype, GBA mutations) as auxiliary input streams; '
        '(3) development of longitudinal prediction models that forecast disease trajectory '
        'rather than binary diagnostic classification; and (4) clinical validation through '
        'prospective studies in memory clinic settings.'
    )

    # --- Page Break: References ---
    doc.add_page_break()
    doc.add_heading('References', level=1)
    references = [
        'Acosta-Cabronero, J., Williams, G. B., Pengas, G., & Nestor, P. J. (2010). Absolute '
        'diffusivities define the landscape of white matter degeneration in Alzheimer\'s disease. '
        'Brain, 133(2), 529-539.',
        'Braak, H., & Braak, E. (1991). Neuropathological staging of Alzheimer-related changes. '
        'Acta Neuropathologica, 82(4), 239-259.',
        'Braak, H., Del Tredici, K., Rub, U., de Vos, R. A., Jansen Steur, E. N., & Braak, E. '
        '(2003). Staging of brain pathology related to sporadic Parkinson\'s disease. '
        'Neurobiology of Aging, 24(2), 197-211.',
        'Dickerson, B. C., Bakkour, A., Salat, D. H., Feczko, E., Pacheco, J., Greve, D. N., '
        '... & Buckner, R. L. (2009). The cortical signature of Alzheimer\'s disease: regionally '
        'specific cortical thinning relates to symptom severity. Cerebral Cortex, 19(3), 497-510.',
        'Frisoni, G. B., Fox, N. C., Jack, C. R., Scheltens, P., & Thompson, P. M. (2010). '
        'The clinical use of structural MRI in Alzheimer disease. Nature Reviews Neurology, '
        '6(2), 67-77.',
        'Hardy, J. A., & Higgins, G. A. (1992). Alzheimer\'s disease: the amyloid cascade '
        'hypothesis. Science, 256(5054), 184-185.',
        'Jack, C. R., Knopman, D. S., Jagust, W. J., Petersen, R. C., Weiner, M. W., Aisen, '
        'P. S., ... & Trojanowski, J. Q. (2013). Tracking pathophysiological processes in '
        'Alzheimer\'s disease: an updated hypothetical model of dynamic biomarkers. '
        'The Lancet Neurology, 12(2), 207-216.',
        'Liu, M., Zhang, J., Adeli, E., & Shen, D. (2023). Joint classification and regression '
        'via deep multi-task multi-channel learning for Alzheimer\'s disease diagnosis. '
        'IEEE Transactions on Biomedical Engineering, 70(4), 1137-1148.',
        'Palmqvist, S., Scholl, M., Strandberg, O., Mattsson, N., Stomrud, E., Zetterberg, H., '
        '... & Hansson, O. (2017). Earliest accumulation of beta-amyloid occurs within the '
        'default-mode network. Nature Communications, 8(1), 1214.',
        'Postuma, R. B., Aarsland, D., Barone, P., Burn, D. J., Hawkes, C. H., Oertel, W., '
        '& Ziemssen, T. (2012). Identifying prodromal Parkinson\'s disease: pre-motor disorders '
        'in Parkinson\'s disease. Movement Disorders, 27(5), 617-626.',
        'Spillantini, M. G., Schmidt, M. L., Lee, V. M., Trojanowski, J. Q., Jakes, R., & '
        'Goedert, M. (1997). Alpha-synuclein in Lewy bodies. Nature, 388(6645), 839-840.',
        'Vaillancourt, D. E., Spraker, M. B., Prodoehl, J., Abraham, I., Corcos, D. M., Zhou, '
        'X. J., ... & Little, D. M. (2009). High-resolution diffusion tensor imaging in the '
        'substantia nigra of de novo Parkinson disease. Neurology, 72(16), 1378-1384.',
        'Wimo, A., Seeher, K., Cataldi, R., Cyhlarova, E., Dielemann, J. L., Frisell, O., '
        '... & Guerchet, M. (2024). The worldwide costs of dementia in 2019. '
        'Alzheimer\'s & Dementia, 19(7), 2865-2873.',
        'World Health Organization. (2023). Dementia Fact Sheet. WHO.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)

    # --- Header & Footer ---
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = hp.add_run('Vasquez-Rodriguez | Machine Learning for Neurodegenerative Disease Detection')
    run.font.size = Pt(9)
    run.italic = True

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fp.text = "Page "
    from docx.oxml.ns import qn
    r1 = fp.add_run()
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
