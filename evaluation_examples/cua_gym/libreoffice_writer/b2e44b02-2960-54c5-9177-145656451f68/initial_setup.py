"""
Initial Setup: Insert cross-reference to table caption in financial report
Task ID: writer_struct_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_072'
OUTPUT = f'/home/user/Desktop/revenue_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark around the paragraph content."""
    # Create bookmarkStart
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), '1')
    bm_start.set(qn('w:name'), bookmark_name)

    # Create bookmarkEnd
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), '1')

    # Insert at beginning and end of paragraph
    paragraph._p.insert(0, bm_start)
    paragraph._p.append(bm_end)


def add_page_break(doc):
    """Add a manual page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return para


def create_initial():
    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # =====================
    # PAGE 1 — Executive Summary
    # =====================
    title_para = doc.add_heading('Annual Financial Report 2024', level=0)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Prepared by: Finance Department | Confidential')
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()  # spacing

    exec_heading = doc.add_heading('Executive Summary', level=1)

    p1 = doc.add_paragraph(
        'This report presents the comprehensive financial performance of Northgate Industries '
        'for the fiscal year ending December 31, 2024. The company achieved record revenues '
        'of $2.4 billion, representing an 18% increase over the prior year. Operating margins '
        'improved to 24.3%, driven by operational efficiency initiatives and favorable market '
        'conditions across all business segments.'
    )

    p2 = doc.add_paragraph(
        'Key performance indicators demonstrate sustained growth momentum. Net income rose '
        'to $583 million, up from $467 million in fiscal year 2023. Earnings per share '
        'increased to $4.72, exceeding analyst consensus estimates of $4.45. The company '
        'continued its strategic expansion into emerging markets, with international revenue '
        'contributing 31% of total revenue, compared to 26% in the prior year.'
    )

    p3 = doc.add_paragraph(
        'Capital expenditures totaled $312 million, primarily directed toward manufacturing '
        'capacity expansion and technology infrastructure upgrades. Free cash flow generation '
        'remained robust at $748 million, enabling the board to approve a 15% dividend '
        'increase and a $500 million share repurchase program.'
    )

    # This is the key paragraph with the incomplete sentence — NO cross-reference here
    p4 = doc.add_paragraph()
    run_p4 = p4.add_run('Revenue data is summarized in ')
    # NOTE: No cross-reference field here — that is what the agent must insert

    doc.add_paragraph()  # spacing

    overview_heading = doc.add_heading('Overview of Key Financial Metrics', level=1)

    metrics_para = doc.add_paragraph(
        'The following sections provide detailed analysis of revenue streams, cost structures, '
        'profitability metrics, and balance sheet position. Each section includes comparative '
        'data from previous fiscal years to illustrate trends and performance trajectories.'
    )

    # =====================
    # PAGE 2 — Revenue Analysis
    # =====================
    add_page_break(doc)

    doc.add_heading('Revenue Analysis', level=1)

    doc.add_paragraph(
        'Total revenue for fiscal year 2024 reached $2,412.6 million, marking the fifth '
        'consecutive year of double-digit revenue growth. The company\'s diversified revenue '
        'base provided resilience against sector-specific headwinds while capitalizing on '
        'growth opportunities across multiple market segments.'
    )

    doc.add_heading('Revenue by Business Segment', level=2)

    doc.add_paragraph(
        'The Technology Solutions segment generated $987.4 million in revenue, representing '
        '40.9% of total revenue and a 22% increase year-over-year. Growth was driven by '
        'strong demand for enterprise software solutions and professional services engagements. '
        'The segment secured 47 new enterprise clients during the year, with an average '
        'contract value of $8.3 million.'
    )

    doc.add_paragraph(
        'Healthcare Services contributed $724.1 million, or 30.0% of total revenue, '
        'reflecting a 14% increase from the prior year. Organic growth was supplemented '
        'by the acquisition of MedTech Innovations in Q2 2024, which added approximately '
        '$85 million in annualized revenue. The segment expanded its service offerings '
        'to include advanced diagnostic imaging and telehealth platforms.'
    )

    doc.add_paragraph(
        'Industrial Manufacturing revenue of $510.8 million represented 21.2% of total '
        'revenue, growing 12% year-over-year. The segment benefited from infrastructure '
        'spending programs and increased automation adoption across manufacturing industries. '
        'Backlog at year-end stood at $1.2 billion, providing strong revenue visibility '
        'for fiscal year 2025.'
    )

    doc.add_paragraph(
        'The Energy & Resources segment generated $190.3 million (7.9% of total revenue), '
        'recovering from a difficult prior year driven by commodity price volatility. '
        'Strategic partnerships with three major energy companies established during the '
        'year are expected to contribute significantly in fiscal 2025.'
    )

    # =====================
    # PAGE 3 — Cost Structure and Profitability
    # =====================
    add_page_break(doc)

    doc.add_heading('Cost Structure and Profitability', level=1)

    doc.add_paragraph(
        'Total operating expenses for fiscal 2024 were $1,826.4 million, representing '
        '75.7% of revenue, compared to 77.2% in the prior year. The improvement in '
        'operating leverage reflects the benefits of scale and ongoing cost optimization '
        'initiatives implemented across the organization.'
    )

    doc.add_heading('Cost of Revenue', level=2)

    doc.add_paragraph(
        'Cost of revenue was $1,124.7 million (46.6% of revenue), compared to $989.3 '
        'million (48.4% of revenue) in fiscal 2023. The improvement reflects a favorable '
        'shift in product mix toward higher-margin software and services offerings, as '
        'well as operational efficiency gains from automation investments made in prior years.'
    )

    doc.add_heading('Operating Expenses', level=2)

    doc.add_paragraph(
        'Research and development expenses totaled $189.4 million (7.9% of revenue), '
        'reflecting continued investment in product innovation and technology advancement. '
        'Key development initiatives include next-generation AI-powered analytics platform, '
        'enhanced cybersecurity solutions, and expanded healthcare interoperability tools.'
    )

    doc.add_paragraph(
        'Selling, general and administrative expenses were $312.3 million (12.9% of '
        'revenue), down from 14.1% in the prior year. Efficiency improvements in sales '
        'operations and corporate functions drove the favorable variance, partially offset '
        'by increased investment in brand building and market development activities.'
    )

    doc.add_paragraph(
        'Operating income for fiscal 2024 was $586.2 million, yielding an operating '
        'margin of 24.3%, compared to 22.8% in fiscal 2023. This improvement of 150 '
        'basis points reflects the cumulative impact of strategic initiatives focused '
        'on revenue quality, operational efficiency, and disciplined cost management.'
    )

    # =====================
    # PAGE 4 — Quarterly Performance (contains the referenced table)
    # =====================
    add_page_break(doc)

    doc.add_heading('Quarterly Performance Review', level=1)

    doc.add_paragraph(
        'The company delivered consistent performance throughout fiscal 2024, with each '
        'quarter contributing to the full-year results. Seasonal patterns typical of the '
        'technology and healthcare sectors were evident, with stronger performance in Q2 '
        'and Q4 driven by enterprise purchasing cycles and year-end budget utilization.'
    )

    # Table captions 1 and 2 (before the referenced table)
    caption1_para = doc.add_paragraph()
    run_c1 = caption1_para.add_run('Table 1: Annual Revenue by Business Segment ($ millions)')
    run_c1.font.bold = True
    run_c1.font.italic = True

    # Table 1
    t1 = doc.add_table(rows=6, cols=3)
    t1.style = 'Table Grid'
    t1_data = [
        ['Business Segment', '2024 Revenue', '2023 Revenue'],
        ['Technology Solutions', '$987.4', '$809.3'],
        ['Healthcare Services', '$724.1', '$635.2'],
        ['Industrial Manufacturing', '$510.8', '$456.1'],
        ['Energy & Resources', '$190.3', '$161.0'],
        ['Total', '$2,412.6', '$2,061.6'],
    ]
    for i, row_data in enumerate(t1_data):
        for j, val in enumerate(row_data):
            cell = t1.cell(i, j)
            cell.text = val
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()  # spacing

    caption2_para = doc.add_paragraph()
    run_c2 = caption2_para.add_run('Table 2: Operating Expenses Summary ($ millions)')
    run_c2.font.bold = True
    run_c2.font.italic = True

    # Table 2
    t2 = doc.add_table(rows=5, cols=3)
    t2.style = 'Table Grid'
    t2_data = [
        ['Expense Category', '2024', '2023'],
        ['Cost of Revenue', '$1,124.7', '$989.3'],
        ['Research & Development', '$189.4', '$161.8'],
        ['Selling, General & Admin', '$312.3', '$290.5'],
        ['Total Operating Expenses', '$1,826.4', '$1,591.3'],
    ]
    for i, row_data in enumerate(t2_data):
        for j, val in enumerate(row_data):
            cell = t2.cell(i, j)
            cell.text = val
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()  # spacing

    # THE KEY TABLE CAPTION — "Table 3: Quarterly Revenue Summary"
    # This caption has a bookmark named "Table_3_Quarterly_Revenue_Summary"
    caption3_para = doc.add_paragraph()
    run_c3 = caption3_para.add_run('Table 3: Quarterly Revenue Summary')
    run_c3.font.bold = True
    run_c3.font.italic = True

    # Add bookmark around caption paragraph for cross-referencing
    add_bookmark(caption3_para, 'Table_3_Quarterly_Revenue_Summary')

    # Table 3: Quarterly Revenue Summary
    t3 = doc.add_table(rows=6, cols=5)
    t3.style = 'Table Grid'
    t3_data = [
        ['Quarter', 'Revenue ($M)', 'Growth YoY', 'Gross Margin', 'Operating Margin'],
        ['Q1 2024', '$561.3', '+15.2%', '52.1%', '22.8%'],
        ['Q2 2024', '$614.7', '+18.4%', '53.8%', '24.1%'],
        ['Q3 2024', '$594.2', '+17.9%', '53.2%', '23.9%'],
        ['Q4 2024', '$642.4', '+19.6%', '54.7%', '25.8%'],
        ['Full Year', '$2,412.6', '+17.5%', '53.4%', '24.3%'],
    ]
    for i, row_data in enumerate(t3_data):
        for j, val in enumerate(row_data):
            cell = t3.cell(i, j)
            cell.text = val
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()  # spacing

    doc.add_paragraph(
        'Quarter-over-quarter analysis reveals consistent execution against financial targets. '
        'The sequential improvement in operating margins across all four quarters demonstrates '
        'the progressive realization of operational leverage as revenue scale increased.'
    )

    # =====================
    # PAGE 5 — Balance Sheet and Cash Flow
    # =====================
    add_page_break(doc)

    doc.add_heading('Balance Sheet and Cash Flow', level=1)

    doc.add_paragraph(
        'The company maintained a strong financial position throughout fiscal 2024, '
        'with total assets growing to $4,823.6 million, up from $4,124.8 million '
        'at the end of fiscal 2023. The balance sheet reflects disciplined capital '
        'allocation and prudent financial management.'
    )

    doc.add_heading('Liquidity Position', level=2)

    doc.add_paragraph(
        'Cash and cash equivalents at year-end stood at $842.5 million, compared to '
        '$634.2 million at the end of fiscal 2023. The strong cash position provides '
        'ample liquidity to fund organic growth initiatives, pursue strategic acquisitions, '
        'and return capital to shareholders.'
    )

    doc.add_paragraph(
        'The company maintains a $1.0 billion revolving credit facility, of which '
        '$750 million remained undrawn at year-end. Total debt was $1,245.3 million, '
        'representing a net debt to EBITDA ratio of 1.4x, well within the company\'s '
        'target leverage range of 1.0x to 2.0x.'
    )

    doc.add_heading('Cash Flow Generation', level=2)

    doc.add_paragraph(
        'Operating cash flow for fiscal 2024 was $834.7 million (34.6% of revenue), '
        'compared to $712.4 million in the prior year. The increase reflects higher '
        'net income and favorable working capital movements, partially offset by '
        'higher tax payments related to the strong operating performance.'
    )

    doc.add_paragraph(
        'Capital expenditures of $86.7 million were directed primarily toward '
        'technology infrastructure, facilities improvements, and manufacturing '
        'equipment upgrades. Free cash flow of $748.0 million represented a '
        'conversion rate of 128% relative to net income, reflecting the asset-light '
        'nature of the company\'s business model.'
    )

    # =====================
    # PAGE 6 — Outlook and Strategy
    # =====================
    add_page_break(doc)

    doc.add_heading('Strategic Outlook and Fiscal 2025 Guidance', level=1)

    doc.add_paragraph(
        'Management is confident in the company\'s ability to sustain its growth '
        'trajectory into fiscal 2025 and beyond. The strategic roadmap focuses on '
        'three key priorities: accelerating organic revenue growth, expanding margins '
        'through operational excellence, and deploying capital in value-enhancing '
        'acquisitions and shareholder returns.'
    )

    doc.add_heading('Fiscal 2025 Financial Guidance', level=2)

    doc.add_paragraph(
        'For fiscal year 2025, management provides the following financial guidance: '
        'total revenue is expected to be in the range of $2,750 million to $2,850 '
        'million, representing growth of 14% to 18% over fiscal 2024. Operating '
        'margin is expected to be in the range of 25.0% to 26.0%, reflecting '
        'continued operational leverage and efficiency improvements.'
    )

    doc.add_paragraph(
        'Earnings per share guidance is set at $5.35 to $5.65, implying growth of '
        '13% to 20% over fiscal 2024 EPS of $4.72. Capital expenditures are '
        'expected to be approximately $100 million, including investments in '
        'capacity expansion and technology modernization programs.'
    )

    doc.add_heading('Strategic Initiatives', level=2)

    doc.add_paragraph(
        'The company will continue to invest in its AI and machine learning capabilities, '
        'building on the foundation established in fiscal 2024. The next-generation '
        'analytics platform is expected to launch commercially in Q2 2025, with initial '
        'deployments at 12 enterprise customers already secured under early adopter programs.'
    )

    doc.add_paragraph(
        'International expansion remains a priority, with targeted investments in '
        'European and Asia-Pacific markets. Management expects international revenue '
        'to grow to approximately 35% of total revenue by fiscal year 2026, supported '
        'by strategic partnerships and selective acquisitions in key growth markets.'
    )

    closing_para = doc.add_paragraph(
        'The board of directors and management team remain committed to delivering '
        'superior long-term value for all stakeholders. The company\'s strong financial '
        'position, differentiated capabilities, and experienced leadership team provide '
        'a solid foundation for continued growth and value creation.'
    )

    # Ensure Desktop directory exists and save
    os.makedirs('/home/user/Desktop', exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
