"""
Reward Script: Insert cross-reference to 'Table 3: Quarterly Revenue Summary'
Task ID: writer_struct_072
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): A REF cross-reference field (fldChar + instrText) exists in paragraph 7
  Component 2 (0.3): Field instruction references bookmark Table_3_Quarterly_Revenue_Summary
  Component 3 (0.2): Field display text contains 'Table 3: Quarterly Revenue Summary'
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'revenue_report'
FILE_PATH = f'{WORKDIR}/{TASK_ID}.docx'

# Word namespace constant
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def w(tag):
    """Return fully-qualified Word XML tag name."""
    return f'{{{W_NS}}}{tag}'


def get_field_display_text(para_xml):
    """
    Extract the display text from a Word field code in the paragraph.
    Locates the 'separate' and 'end' fldChar marker indices, then collects
    text from runs between them. Returns combined text or empty string.
    """
    all_runs = para_xml.findall(f'.//{w("r")}')
    separate_idx = None
    end_idx = None
    for idx, run in enumerate(all_runs):
        fc = run.find(w('fldChar'))
        if fc is not None:
            ft = fc.get(w('fldCharType'))
            if ft == 'separate' and separate_idx is None:
                separate_idx = idx
            elif ft == 'end' and separate_idx is not None and end_idx is None:
                end_idx = idx
    if separate_idx is None or end_idx is None:
        return ''
    parts = []
    for run in all_runs[separate_idx + 1:end_idx]:
        t_elem = run.find(w('t'))
        if t_elem is not None and t_elem.text:
            parts.append(t_elem.text)
    return ''.join(parts)


def verify_task(file_path):
    """
    Verify that a cross-reference field was inserted in paragraph 7.
    The task requires inserting a REF field after 'Revenue data is summarized in '
    that references the bookmark 'Table_3_Quarterly_Revenue_Summary' and displays
    'Table 3: Quarterly Revenue Summary'.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Sanity check: document has expected structure
    if len(doc.paragraphs) < 8:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 8")
        print("REWARD: 0.0")
        return 0.0

    # Target paragraph: index 7 — "Revenue data is summarized in <field>"
    para = doc.paragraphs[7]
    para_xml = para._element

    # -------------------------------------------------------------------------
    # Component 1: A fldChar field exists in paragraph 7 (0.5 points)
    # Checks: fldChar elements with fldCharType="begin" AND instrText elements present
    # FAILS on initial (no field) → PASSES on golden (field present)
    # -------------------------------------------------------------------------
    try:
        fld_chars = para_xml.findall(f'.//{w("fldChar")}')
        begin_count = sum(
            1 for fc in fld_chars if fc.get(w('fldCharType')) == 'begin'
        )
        instr_count = len(para_xml.findall(f'.//{w("instrText")}'))

        if begin_count > 0 and instr_count > 0:
            print(f"PASS: Component 1 — Cross-reference field (fldChar + instrText) found in paragraph 7 (0.5 pts)")
            total_score += 0.5
        else:
            print(
                f"FAIL: Component 1 — No field code found in paragraph 7 "
                f"(fldChar begin={begin_count}, instrText={instr_count})"
            )
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Field instruction references the correct bookmark (0.3 points)
    # Verifies instrText contains both 'REF' and 'Table_3_Quarterly_Revenue_Summary'
    # FAILS on initial (no instrText) → PASSES on golden (correct REF instruction)
    # -------------------------------------------------------------------------
    try:
        instr_elems = para_xml.findall(f'.//{w("instrText")}')
        matching_instrs = [
            instr for instr in instr_elems
            if instr.text and 'REF' in instr.text and 'Table_3_Quarterly_Revenue_Summary' in instr.text
        ]
        if len(matching_instrs) > 0:
            print(f"PASS: Component 2 — instrText references 'Table_3_Quarterly_Revenue_Summary' (0.3 pts)")
            print(f"       Field instruction: {matching_instrs[0].text.strip()!r}")
            total_score += 0.3
        else:
            all_instr_values = [i.text for i in instr_elems]
            print(
                f"FAIL: Component 2 — No instrText with 'REF Table_3_Quarterly_Revenue_Summary' found. "
                f"Found: {all_instr_values}"
            )
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Field display text contains the caption text (0.2 points)
    # Checks text in runs between 'separate' and 'end' fldChar markers
    # FAILS on initial (no field runs) → PASSES on golden (correct display text)
    # -------------------------------------------------------------------------
    try:
        display_text = get_field_display_text(para_xml)
        expected_caption = 'Table 3: Quarterly Revenue Summary'
        if expected_caption in display_text:
            print(f"PASS: Component 3 — Field display text contains '{expected_caption}' (0.2 pts)")
            total_score += 0.2
        else:
            print(
                f"FAIL: Component 3 — Expected display text '{expected_caption}', "
                f"found: {display_text!r}"
            )
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
