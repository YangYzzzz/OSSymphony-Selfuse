"""
Reward Script: Create a multi-level outline list for project plan
Task ID: writer_list_005
Domain: libreoffice_writer
Scoring:
  Component 1: Level-1 items (Research Phase, Design Phase, Implementation Phase) have
               numbering at ilvl=0 in a multi-level list — 0.35 pts
  Component 2: Level-2 sub-items have numbering at ilvl=1 in the same multi-level list
               — 0.35 pts
  Component 3: Numbering format is correct (%1. for level-0, %1.%2 for level-1)
               — 0.20 pts
  Component 4: Level-2 paragraphs are indented further than level-1 paragraphs
               — 0.10 pts
Total: 1.0
"""

import subprocess
subprocess.run(['pip3', 'install', 'python-docx'], check=True, capture_output=True)

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_list_005'
FILE_PATH = f'{WORKDIR}/Desktop/project_plan.docx'

# The expected paragraphs in order (text, expected_ilvl)
EXPECTED_OUTLINE = [
    ('Research Phase', 0),
    ('Conduct market analysis', 1),
    ('Review competitor products', 1),
    ('Design Phase', 0),
    ('Create wireframes', 1),
    ('Design user interface mockups', 1),
    ('Write technical specifications', 1),
    ('Implementation Phase', 0),
    ('Set up development environment', 1),
    ('Develop core features', 1),
]

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_para_numpr(para):
    """Return (ilvl, numId) from paragraph numPr element, or (None, None) if absent."""
    pPr = para._element.find(f'{{{W_NS}}}pPr')
    if pPr is None:
        return None, None
    numPr = pPr.find(f'{{{W_NS}}}numPr')
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(f'{{{W_NS}}}ilvl')
    numId_el = numPr.find(f'{{{W_NS}}}numId')
    ilvl = ilvl_el.get(f'{{{W_NS}}}val') if ilvl_el is not None else None
    numId = numId_el.get(f'{{{W_NS}}}val') if numId_el is not None else None
    return ilvl, numId


def get_abstract_num_formats(doc, target_num_id):
    """
    Given a concrete numId, return the lvlText formats for ilvl=0 and ilvl=1.
    Returns (lvlText_ilvl0, lvlText_ilvl1) or (None, None) if not found.
    """
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return None, None

    xml_root = numbering_part._element

    # Step 1: find abstractNumId for this numId
    abstract_num_id_val = None
    for num_el in xml_root.findall(f'{{{W_NS}}}num'):
        if num_el.get(f'{{{W_NS}}}numId') == str(target_num_id):
            abstract_ref = num_el.find(f'{{{W_NS}}}abstractNumId')
            if abstract_ref is not None:
                abstract_num_id_val = abstract_ref.get(f'{{{W_NS}}}val')
            break

    if abstract_num_id_val is None:
        return None, None

    # Step 2: find abstractNum with that abstractNumId
    lvl_text = {}
    lvl_left_indent = {}
    for abstract_el in xml_root.findall(f'{{{W_NS}}}abstractNum'):
        if abstract_el.get(f'{{{W_NS}}}abstractNumId') == abstract_num_id_val:
            for lvl_el in abstract_el.findall(f'{{{W_NS}}}lvl'):
                ilvl = lvl_el.get(f'{{{W_NS}}}ilvl')
                lvl_text_el = lvl_el.find(f'{{{W_NS}}}lvlText')
                if lvl_text_el is not None:
                    lvl_text[ilvl] = lvl_text_el.get(f'{{{W_NS}}}val', '')
                # Check indentation
                pPr_el = lvl_el.find(f'{{{W_NS}}}pPr')
                if pPr_el is not None:
                    ind_el = pPr_el.find(f'{{{W_NS}}}ind')
                    if ind_el is not None:
                        left_val = ind_el.get(f'{{{W_NS}}}left')
                        if left_val is not None:
                            lvl_left_indent[ilvl] = int(left_val)
            break

    fmt0 = lvl_text.get('0', None)
    fmt1 = lvl_text.get('1', None)
    ind0 = lvl_left_indent.get('0', None)
    ind1 = lvl_left_indent.get('1', None)
    return fmt0, fmt1, ind0, ind1


def verify_task(file_path):
    """
    Verify the multi-level outline list in project_plan.docx.
    Returns a float in [0.0, 1.0].
    """
    total_score = 0.0

    # Precondition gate: file must exist and be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build paragraph map: text -> (ilvl, numId)
    para_info = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            ilvl, numId = get_para_numpr(para)
            para_info[text] = (ilvl, numId)

    # -------------------------------------------------------
    # Component 1: Level-1 items (ilvl=0) have numbering (0.35 pts)
    # -------------------------------------------------------
    level1_items = [('Research Phase', 0), ('Design Phase', 0), ('Implementation Phase', 0)]
    level1_pass = 0
    level1_numId = None

    try:
        for (text, expected_ilvl) in level1_items:
            if text in para_info:
                ilvl, numId = para_info[text]
                if ilvl == str(expected_ilvl) and numId is not None and numId != '0':
                    level1_pass += 1
                    if level1_numId is None:
                        level1_numId = numId
                    print(f"PASS: '{text}' has ilvl={ilvl}, numId={numId}")
                else:
                    print(f"FAIL: '{text}' expected ilvl=0 with numId, found ilvl={ilvl}, numId={numId}")
            else:
                print(f"FAIL: '{text}' not found in document paragraphs")

        if level1_pass == len(level1_items):
            print(f"PASS: Component 1 — all 3 level-1 items numbered at ilvl=0 (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — only {level1_pass}/{len(level1_items)} level-1 items numbered correctly")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------
    # Component 2: Level-2 items (ilvl=1) have numbering (0.35 pts)
    # -------------------------------------------------------
    level2_items = [
        'Conduct market analysis',
        'Review competitor products',
        'Create wireframes',
        'Design user interface mockups',
        'Write technical specifications',
        'Set up development environment',
        'Develop core features',
    ]
    level2_pass = 0
    level2_numId = None

    try:
        for text in level2_items:
            if text in para_info:
                ilvl, numId = para_info[text]
                if ilvl == '1' and numId is not None and numId != '0':
                    level2_pass += 1
                    if level2_numId is None:
                        level2_numId = numId
                    print(f"PASS: '{text}' has ilvl={ilvl}, numId={numId}")
                else:
                    print(f"FAIL: '{text}' expected ilvl=1 with numId, found ilvl={ilvl}, numId={numId}")
            else:
                print(f"FAIL: '{text}' not found in document paragraphs")

        if level2_pass == len(level2_items):
            print(f"PASS: Component 2 — all 7 level-2 items numbered at ilvl=1 (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 2 — only {level2_pass}/{len(level2_items)} level-2 items numbered correctly")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------
    # Component 3: Numbering format is correct (%1. for ilvl=0, %1.%2 for ilvl=1) (0.20 pts)
    # -------------------------------------------------------
    try:
        # Identify the numId used for the list (prefer level1_numId or level2_numId)
        primary_numId = level1_numId or level2_numId
        if primary_numId is None:
            print("FAIL: Component 3 — no numId found, cannot check format")
        else:
            fmt0, fmt1, ind0, ind1 = get_abstract_num_formats(doc, int(primary_numId))
            print(f"INFO: abstractNum for numId={primary_numId}: lvlText[0]={repr(fmt0)}, lvlText[1]={repr(fmt1)}")
            print(f"INFO: left indent: ilvl=0 -> {ind0}, ilvl=1 -> {ind1}")

            # Level-0 should be "%1." format (decimal with period)
            fmt0_ok = fmt0 is not None and '%1' in fmt0 and fmt0.endswith('.')
            # Level-1 should be "%1.%2" format (parent.child decimal)
            fmt1_ok = fmt1 is not None and '%1' in fmt1 and '%2' in fmt1

            if fmt0_ok and fmt1_ok:
                print(f"PASS: Component 3 — numbering formats correct: lvl0='{fmt0}', lvl1='{fmt1}' (0.20 pts)")
                total_score += 0.20
            else:
                if not fmt0_ok:
                    print(f"FAIL: Component 3 — level-0 format '{fmt0}' does not match expected '%1.' pattern")
                if not fmt1_ok:
                    print(f"FAIL: Component 3 — level-1 format '{fmt1}' does not match expected '%1.%2' pattern")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------
    # Component 4: Level-2 items are indented further than level-1 items (0.10 pts)
    # -------------------------------------------------------
    try:
        primary_numId_for_indent = level1_numId or level2_numId
        if primary_numId_for_indent is None:
            print("FAIL: Component 4 — no numId found, cannot check indentation")
        else:
            fmt0_i, fmt1_i, ind0, ind1 = get_abstract_num_formats(doc, int(primary_numId_for_indent))
            if ind0 is not None and ind1 is not None:
                if ind1 > ind0:
                    print(f"PASS: Component 4 — level-2 indent ({ind1}) > level-1 indent ({ind0}) (0.10 pts)")
                    total_score += 0.10
                else:
                    print(f"FAIL: Component 4 — level-2 indent ({ind1}) NOT greater than level-1 indent ({ind0})")
            else:
                # Fall back to checking paragraph-level indentation
                # Some implementations use paragraph pPr instead of abstractNum pPr
                l1_indents = []
                l2_indents = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    ilvl, numId = get_para_numpr(para)
                    pf = para.paragraph_format
                    left_ind = pf.left_indent
                    if ilvl == '0' and numId is not None:
                        l1_indents.append(left_ind if left_ind is not None else 0)
                    elif ilvl == '1' and numId is not None:
                        l2_indents.append(left_ind if left_ind is not None else 0)

                if l1_indents and l2_indents:
                    avg_l1 = sum(l1_indents) / len(l1_indents)
                    avg_l2 = sum(l2_indents) / len(l2_indents)
                    if avg_l2 > avg_l1:
                        print(f"PASS: Component 4 — para-level: avg l2 indent ({avg_l2}) > avg l1 indent ({avg_l1}) (0.10 pts)")
                        total_score += 0.10
                    else:
                        print(f"FAIL: Component 4 — para-level: avg l2 indent ({avg_l2}) NOT > avg l1 indent ({avg_l1})")
                else:
                    print(f"FAIL: Component 4 — could not determine indentation (ind0={ind0}, ind1={ind1})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
