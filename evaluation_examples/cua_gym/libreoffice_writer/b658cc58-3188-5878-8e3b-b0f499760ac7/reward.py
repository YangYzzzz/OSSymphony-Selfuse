"""
FINAL REWARD SCRIPT - SUCCESS
Task: Paragraph 4 is the odd one out—everything else looks fine, but that single paragraph still sits in plain sentence case. In LibreOffice Writer, how do I convert just paragraph 4 to Title Case so that every word starts with a capital letter, while leaving the rest of the document untouched?
Generated: 2025-09-10 13:12:13
Status: success
Model: azure-o3
Total Steps: 2
"""

import os
import re
from docx import Document


def is_word_title_case(word: str) -> bool:
    """Return True if the alphabetic portion of `word` is in Title Case.
    Non-alphabetic characters (punctuation) are ignored.
    Example:
        "Hello" -> True
        "WORLD" -> False (all caps)
        "eBook" -> False (starts lower)
    """
    letters_only = re.sub(r"[^A-Za-z]", "", word)
    if not letters_only:  # word contains no letters (e.g., purely punctuation)
        return True  # punctuation doesn't violate title-case requirement

    # First character must be uppercase, the rest must be lowercase.
    return letters_only[0].isupper() and letters_only[1:].islower()


def verify_paragraph4_title_case(file_path: str) -> float:
    """Verify that ONLY paragraph 4 of the document is converted to Title Case.

    Scoring rubric (progressive):
        – 0.1  Document has at least 4 paragraphs with non-trivial text
        – 0.7  Every word in paragraph 4 is correctly Title Cased
        – 0.2  At least one word shows true mixed case (i.e., not ALL CAPS)
    Total possible: 1.0
    """
    max_score = 1.0
    total_score = 0.0

    # ---------- 1. Load document ----------
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0  # Cannot proceed without the file

    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Failed to open DOCX: {exc}")
        print("REWARD: 0.0")
        return 0.0

    para_count = len(doc.paragraphs)
    print(f"Loaded document with {para_count} paragraph(s)")

    # ---------- 2. Basic prerequisite check ----------
    if para_count < 4:
        print("✗ Document has fewer than 4 paragraphs — task impossible")
        print("REWARD: 0.0")
        return 0.0

    # Only paragraph 4 (index 3) is evaluated for Title Case
    paragraph4_text = doc.paragraphs[3].text.strip()
    print(f"Paragraph 4 text: '{paragraph4_text}'")

    # Tokenise on whitespace to evaluate each word
    words = re.split(r"\s+", paragraph4_text)
    alpha_words = [w for w in words if re.search(r"[A-Za-z]", w)]

    # ---------- 3. Scoring begins ----------
    # 3a. Ensure paragraph has substantive content (>=4 alphabetic words)
    if len(alpha_words) >= 4:
        print("✓ Paragraph contains sufficient words (0.1)")
        total_score += 0.1
    else:
        print("✗ Paragraph too short — no points for content length")

    # 3b. Check that every word is Title Cased
    all_title_case = True
    for word in words:
        if not is_word_title_case(word):
            all_title_case = False
            print(f"✗ Word not title case: '{word}'")
            break

    if all_title_case:
        print("✓ All words are Title Cased (0.7)")
        total_score += 0.7

    # 3c. Ensure words are not simply ALL CAPS (needs lowercase after first letter)
    proper_mixed_case = any(
        len(re.sub(r"[^A-Za-z]", "", w)) > 1 and re.sub(r"[^A-Za-z]", "", w)[1:].islower()
        for w in words
    )
    if proper_mixed_case:
        print("✓ Proper mixed case verified (0.2)")
        total_score += 0.2
    else:
        print("✗ Detected potential ALL-CAPS — no points for mixed-case check")

    # ---------- 4. Finalise score ----------
    final_score = min(total_score, max_score)
    print(f"Total score: {final_score}")
    print(f"REWARD: {final_score}")
    return final_score


# -------------------- EXECUTION ENTRY POINT --------------------
if __name__ == "__main__":
    TEST_FILE = "/home/user/paragraph_4_is_the_odd_one_outeverything_else_looks_fine_but_that_single_paragraph_still_sits_in_pla.docx"
    verify_paragraph4_title_case(TEST_FILE)
