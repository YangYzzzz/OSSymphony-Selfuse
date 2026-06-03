"""
Initial Setup: Software user guide with centered page numbers in footer
Task ID: writer_page_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_page_030'
# Context specifies file on Desktop named user_guide.docx
OUTPUT = f'{WORKDIR}/Desktop/user_guide.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_centered(paragraph):
    """Add a centered page number field (PAGE) to a paragraph."""
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # fldChar begin
    r_begin = paragraph.add_run()
    fldChar_begin = r_begin._element.makeelement(
        qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}
    )
    r_begin._element.append(fldChar_begin)
    # instrText
    r_instr = paragraph.add_run()
    instr = r_instr._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r_instr._element.append(instr)
    # fldChar end
    r_end = paragraph.add_run()
    fldChar_end = r_end._element.makeelement(
        qn('w:fldChar'), {qn('w:fldCharType'): 'end'}
    )
    r_end._element.append(fldChar_end)


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # --- Page setup: A4 portrait, all margins 2.54 cm ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ---- Document Content: 10-page software user guide ----

    # Page 1 - Title & Introduction
    doc.add_heading('CloudSync Pro — User Guide', level=0)
    doc.add_paragraph(
        'Welcome to CloudSync Pro, the enterprise file synchronization and '
        'collaboration platform. This guide covers installation, configuration, '
        'day-to-day usage, and advanced administration features.'
    )
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'CloudSync Pro version 4.2 supports Windows 10/11, macOS 12+, and '
        'Ubuntu 22.04 LTS. The system requires a minimum of 4 GB RAM, '
        '2 GHz dual-core processor, and 500 MB disk space for the client software.'
    )
    doc.add_paragraph(
        'Key features include real-time synchronization across up to 25 devices, '
        'end-to-end AES-256 encryption, version history up to 365 days, '
        'and integration with Microsoft 365, Google Workspace, and Slack.'
    )
    doc.add_page_break()

    # Page 2 - System Requirements
    doc.add_heading('Chapter 1: System Requirements', level=1)
    doc.add_paragraph(
        'Before installing CloudSync Pro, verify that your system meets the '
        'following minimum and recommended specifications.'
    )
    doc.add_heading('1.1 Minimum Requirements', level=2)
    for item in [
        'Operating System: Windows 10 (64-bit), macOS 12.0, or Ubuntu 22.04',
        'Processor: 2 GHz dual-core (Intel Core i3 or AMD Ryzen 3)',
        'RAM: 4 GB',
        'Storage: 500 MB available disk space',
        'Network: Broadband Internet connection (5 Mbps or faster)',
        'Screen Resolution: 1280 × 720',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('1.2 Recommended Requirements', level=2)
    for item in [
        'Processor: 2.5 GHz quad-core (Intel Core i5/i7 or AMD Ryzen 5/7)',
        'RAM: 8 GB or more',
        'Storage: 2 GB available disk space (plus sync folder)',
        'Network: 25 Mbps or faster for large team deployments',
        'Screen Resolution: 1920 × 1080',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # Page 3 - Installation
    doc.add_heading('Chapter 2: Installation', level=1)
    doc.add_heading('2.1 Windows Installation', level=2)
    doc.add_paragraph(
        'Download the CloudSync Pro installer (CloudSyncPro_4.2_Setup.exe) from '
        'the official portal at https://portal.cloudsyncpro.com/downloads. '
        'Run the installer with administrator privileges.'
    )
    for i, step in enumerate([
        'Double-click CloudSyncPro_4.2_Setup.exe to launch the installer.',
        'Accept the End User License Agreement and click Next.',
        'Choose the installation directory (default: C:\\Program Files\\CloudSync Pro).',
        'Select optional components: Desktop icon, Quick Launch shortcut.',
        'Click Install. The installation takes approximately 2–4 minutes.',
        'Click Finish. CloudSync Pro launches automatically on first run.',
    ], 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')
    doc.add_heading('2.2 macOS Installation', level=2)
    doc.add_paragraph(
        'Download CloudSyncPro_4.2.dmg from the portal. Open the DMG file, '
        'drag CloudSync Pro to your Applications folder, and launch it from Spotlight.'
    )
    doc.add_heading('2.3 Linux Installation', level=2)
    doc.add_paragraph(
        'For Ubuntu/Debian systems, add the CloudSync Pro APT repository and '
        'install via: sudo apt-get install cloudsyncpro. '
        'For RPM-based systems use the provided .rpm package.'
    )
    doc.add_page_break()

    # Page 4 - Account Setup
    doc.add_heading('Chapter 3: Account Setup and Sign-In', level=1)
    doc.add_heading('3.1 Creating a New Account', level=2)
    doc.add_paragraph(
        'Navigate to https://portal.cloudsyncpro.com/register and complete '
        'the registration form with your business email address. '
        'A verification email will be sent within 2 minutes.'
    )
    doc.add_heading('3.2 Enterprise SSO Configuration', level=2)
    doc.add_paragraph(
        'Enterprise customers can configure Single Sign-On (SSO) with SAML 2.0 '
        'or OpenID Connect. Contact your IT administrator to obtain the IdP '
        'metadata URL and configure it in the Admin Console under '
        'Settings → Authentication → SSO.'
    )
    doc.add_heading('3.3 Two-Factor Authentication', level=2)
    doc.add_paragraph(
        'Two-factor authentication (2FA) is strongly recommended. '
        'Supported methods: authenticator apps (Google Authenticator, Authy), '
        'SMS verification, and hardware security keys (FIDO2/WebAuthn).'
    )
    doc.add_page_break()

    # Page 5 - Sync Setup
    doc.add_heading('Chapter 4: Setting Up Synchronization', level=1)
    doc.add_heading('4.1 Choosing Your Sync Folder', level=2)
    doc.add_paragraph(
        'During first-run setup, you will be prompted to select a local sync '
        'folder. The default location is ~/CloudSync Pro on macOS/Linux and '
        'C:\\Users\\<username>\\CloudSync Pro on Windows. '
        'You can change this at any time in Preferences → Sync.'
    )
    doc.add_heading('4.2 Selective Sync', level=2)
    doc.add_paragraph(
        'Large teams with extensive shared storage may benefit from Selective Sync. '
        'Under Preferences → Account → Selective Sync, choose which remote '
        'folders to mirror locally. Unselected folders remain in the cloud but '
        'are not downloaded to the current device.'
    )
    doc.add_heading('4.3 Bandwidth Throttling', level=2)
    doc.add_paragraph(
        'To prevent CloudSync Pro from consuming all available bandwidth during '
        'business hours, configure upload and download limits under '
        'Preferences → Network → Bandwidth. '
        'Schedule profiles allow different limits for working hours vs. off-hours.'
    )
    doc.add_page_break()

    # Page 6 - File Sharing
    doc.add_heading('Chapter 5: Sharing Files and Folders', level=1)
    doc.add_heading('5.1 Sharing with Team Members', level=2)
    doc.add_paragraph(
        'Right-click any file or folder in the sync directory and select '
        '"Share with CloudSync Pro". Enter the recipient\'s email address or '
        'select from your team directory. Choose permission level: '
        'View Only, Can Comment, or Can Edit.'
    )
    doc.add_heading('5.2 Public Share Links', level=2)
    doc.add_paragraph(
        'Generate a shareable link for external recipients via '
        'right-click → Share → Copy Link. Links can have optional password '
        'protection and expiry dates (1 day, 7 days, 30 days, or never). '
        'Track link access in the Sharing dashboard under Activity → Links.'
    )
    doc.add_heading('5.3 Team Folders', level=2)
    doc.add_paragraph(
        'Admins can create Team Folders accessible to entire departments or '
        'custom groups. Navigate to Admin Console → Team Folders → New. '
        'Assign members and set default permissions. '
        'Team folder changes sync to all member devices within 60 seconds.'
    )
    doc.add_page_break()

    # Page 7 - Version History
    doc.add_heading('Chapter 6: Version History and Recovery', level=1)
    doc.add_heading('6.1 Viewing File History', level=2)
    doc.add_paragraph(
        'Every time a synced file is modified, CloudSync Pro saves a version '
        'snapshot. To view history, right-click the file and select '
        '"Version History". The history panel shows timestamp, file size, '
        'and the user who made the change.'
    )
    doc.add_heading('6.2 Restoring Previous Versions', level=2)
    doc.add_paragraph(
        'In the Version History panel, click any version to preview it. '
        'Select "Restore" to replace the current version, or "Download" to '
        'retrieve a copy without overwriting the current file. '
        'Restoration triggers a sync event visible to all collaborators.'
    )
    doc.add_heading('6.3 Deleted File Recovery', level=2)
    doc.add_paragraph(
        'Files deleted from the sync folder are moved to the CloudSync Pro '
        'Trash, accessible from the web portal. Deleted files are retained '
        'for 30 days (Business plan) or 365 days (Enterprise plan) '
        'before permanent deletion.'
    )
    doc.add_page_break()

    # Page 8 - Security
    doc.add_heading('Chapter 7: Security and Compliance', level=1)
    doc.add_heading('7.1 Encryption', level=2)
    doc.add_paragraph(
        'CloudSync Pro uses TLS 1.3 for data in transit and AES-256 for data '
        'at rest. Enterprise plan subscribers can enable client-side encryption '
        '(CSE), meaning files are encrypted on your device before upload. '
        'With CSE enabled, CloudSync Pro staff cannot access file contents.'
    )
    doc.add_heading('7.2 Compliance Certifications', level=2)
    for cert in [
        'SOC 2 Type II (audit report available upon request)',
        'ISO 27001:2022 certified',
        'GDPR compliant — data residency options: US, EU, APAC',
        'HIPAA Business Associate Agreement available for healthcare customers',
        'FedRAMP Moderate authorization (US government customers)',
    ]:
        doc.add_paragraph(cert, style='List Bullet')
    doc.add_heading('7.3 Audit Logs', level=2)
    doc.add_paragraph(
        'Admins can access comprehensive audit logs in Admin Console → Audit. '
        'Logs include login events, file operations, permission changes, and '
        'admin actions. Export logs in CSV or JSON format for SIEM integration.'
    )
    doc.add_page_break()

    # Page 9 - Troubleshooting
    doc.add_heading('Chapter 8: Troubleshooting', level=1)
    doc.add_heading('8.1 Common Sync Issues', level=2)
    doc.add_paragraph(
        'If files are not syncing, first check the status indicator in the '
        'system tray (green = synced, yellow = syncing, red = error). '
        'Hover over the icon to see a summary of pending operations or errors.'
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Error Code'
    hdr[1].text = 'Description'
    hdr[2].text = 'Resolution'
    errors = [
        ('E1001', 'Authentication failure', 'Re-enter credentials in Preferences → Account'),
        ('E1042', 'File path too long (Windows)', 'Shorten file/folder names; max 260 chars'),
        ('E2003', 'Insufficient cloud storage', 'Upgrade plan or delete unused files'),
        ('E3011', 'Network timeout', 'Check internet connection; disable VPN temporarily'),
        ('E4005', 'Conflicting file lock', 'Close file in all applications before syncing'),
    ]
    for code, desc, res in errors:
        row = table.add_row().cells
        row[0].text = code
        row[1].text = desc
        row[2].text = res
    doc.add_paragraph('')
    doc.add_heading('8.2 Contacting Support', level=2)
    doc.add_paragraph(
        'If the issue persists, collect diagnostic logs via Help → Collect Logs '
        'and submit them to support@cloudsyncpro.com with a description of the '
        'problem. Priority support is available 24/7 for Enterprise customers '
        'via the dedicated Slack channel or phone line +1-800-CSYNC-01.'
    )
    doc.add_page_break()

    # Page 10 - Appendix
    doc.add_heading('Appendix: Keyboard Shortcuts and Glossary', level=1)
    doc.add_heading('A.1 Keyboard Shortcuts', level=2)
    shortcuts = [
        ('Ctrl+Shift+S (Win/Linux) / Cmd+Shift+S (Mac)', 'Force sync now'),
        ('Ctrl+Shift+P / Cmd+Shift+P', 'Open Preferences'),
        ('Ctrl+Shift+L / Cmd+Shift+L', 'Open Activity Log'),
        ('Ctrl+Z / Cmd+Z', 'Undo last sync operation (where supported)'),
        ('F5 (Windows)', 'Refresh folder view'),
    ]
    for shortcut, action in shortcuts:
        doc.add_paragraph(f'{shortcut}  —  {action}', style='List Bullet')
    doc.add_heading('A.2 Glossary', level=2)
    glossary = [
        ('CSE', 'Client-Side Encryption: files encrypted locally before upload.'),
        ('Delta Sync', 'Only changed file blocks are transferred, reducing bandwidth.'),
        ('IdP', 'Identity Provider: the SSO authentication server (e.g., Okta, Azure AD).'),
        ('Selective Sync', 'Choosing which cloud folders to mirror on a specific device.'),
        ('Version Snapshot', 'A point-in-time copy of a file saved automatically on each change.'),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        run_term = p.add_run(f'{term}: ')
        run_term.bold = True
        p.add_run(definition)

    # --- Footer: page numbers centered ---
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear any existing paragraphs
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ''
    fp = footer.paragraphs[0]
    add_page_number_centered(fp)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
