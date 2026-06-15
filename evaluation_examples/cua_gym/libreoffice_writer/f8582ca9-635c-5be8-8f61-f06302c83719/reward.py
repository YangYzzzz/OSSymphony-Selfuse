"""
Reward Script: Create a TOC at the beginning of the document with title 'Document Index',
including levels 1 through 3, and configure the TOC background to light gray.
Task ID: writer_struct_068
Domain: libreoffice_writer
Scoring:
  Component 1: TOC heading exists at start with title 'Document Index' (0.35 pts)
  Component 2: TOC field covers levels 1-3 via instrText field code (0.35 pts)
  Component 3: Light gray background configured on TOC (0.30 pts)
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_068'
FILE_PATH = '/home/user/Desktop/compliance_manual.docx'
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print('CRITICAL: Cannot load file %s: %s' % (file_path, e))
        print('REWARD: 0.0')
        return 0.0

    body = doc.element.body

    # Component 1: TOC heading at document start with title 'Document Index' (0.35 points)
    # The first paragraph must have a TOC Heading style and text exactly 'Document Index'.
    # This FAILS on initial (no TOC heading exists) and PASSES on golden.
    try:
        first_para = doc.paragraphs[0] if doc.paragraphs else None
        has_toc_heading_style = (
            first_para is not None
            and ('TOC' in first_para.style.name or 'toc' in first_para.style.name.lower())
        )
        has_correct_title = (
            first_para is not None
            and first_para.text.strip() == 'Document Index'
        )

        if has_toc_heading_style and has_correct_title:
            print('PASS: Component 1 — TOC heading at document start with title "Document Index" (0.35 pts)')
            total_score += 0.35
        else:
            details = []
            if not has_toc_heading_style:
                first_style = first_para.style.name if first_para else 'N/A'
                details.append('first paragraph style is %r, not TOC Heading' % first_style)
            if not has_correct_title:
                first_text = first_para.text.strip() if first_para else ''
                details.append('title is %r, expected "Document Index"' % first_text)
            print('FAIL: Component 1 — %s' % '; '.join(details))
    except Exception as e:
        print('ERROR: Component 1 — %s' % e)

    # Component 2: TOC field code references levels 1-3 (0.35 points)
    # The TOC should have instrText containing TOC field with \o "1-3".
    # This FAILS on initial (no instrText at all) and PASSES on golden.
    try:
        instr_elements = body.findall('.//{%s}instrText' % NS)
        toc_instr_texts = [el.text for el in instr_elements if el.text and 'TOC' in el.text]

        # Check if any TOC instruction covers levels 1-3 (from_level <= 1, to_level >= 3)
        level_match = None
        for instr_text in toc_instr_texts:
            m = re.search(r'\\o\s+"(\d+)-(\d+)"', instr_text)
            if m and int(m.group(1)) <= 1 and int(m.group(2)) >= 3:
                level_match = instr_text.strip()
                break

        if level_match is not None:
            print('PASS: Component 2 — TOC field with levels 1-3 found (instrText: %r) (0.35 pts)' % level_match)
            total_score += 0.35
        elif toc_instr_texts:
            print('FAIL: Component 2 — TOC field found but no 1-3 level range. instrText: %r' % toc_instr_texts)
        else:
            sdt_count = len(body.findall('.//{%s}sdt' % NS))
            print('FAIL: Component 2 — No TOC field instruction found (SDT count=%d, instrText count=%d)' % (sdt_count, len(instr_elements)))
    except Exception as e:
        print('ERROR: Component 2 — %s' % e)

    # Component 3: Light gray background on TOC elements (0.30 points)
    # Shading fill D3D3D3 (or similar light gray) must be present on TOC heading/content.
    # This FAILS on initial (no shading exists at all) and PASSES on golden.
    try:
        shd_elements = body.findall('.//{%s}shd' % NS)
        found_fills = [
            shd.get('{%s}fill' % NS).upper()
            for shd in shd_elements
            if shd.get('{%s}fill' % NS)
        ]

        # Check for light gray: D3D3D3 or any near-white gray (R,G,B all > 180, close together)
        gray_fills = [
            f for f in found_fills
            if len(f) == 6
            and all(int(f[i*2:(i+1)*2], 16) > 180 for i in range(3))
            and max(int(f[i*2:(i+1)*2], 16) for i in range(3)) - min(int(f[i*2:(i+1)*2], 16) for i in range(3)) < 30
        ]

        if gray_fills:
            print('PASS: Component 3 — Light gray background found on TOC (fills=%r) (0.30 pts)' % gray_fills)
            total_score += 0.30
        else:
            print('FAIL: Component 3 — No light gray background found on TOC. All fills: %r' % found_fills[:5])
    except Exception as e:
        print('ERROR: Component 3 — %s' % e)

    final_score = min(total_score, 1.0)
    print('\nScore: %.2f/1.0' % total_score)
    print('REWARD: %.1f' % final_score)
    return final_score


if not os.path.exists(FILE_PATH):
    print('File not found: %s' % FILE_PATH)
    print('REWARD: 0.0')
else:
    verify_task(FILE_PATH)
