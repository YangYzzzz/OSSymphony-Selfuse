"""
Initial Setup: Compliance manual document without TOC for writer_struct_068
Task ID: writer_struct_068
Domain: libreoffice_writer

Creates a 20-page compliance manual with 8 Heading 1, 20 Heading 2, and 15 Heading 3 entries,
saved to /home/user/Desktop/compliance_manual.docx. No TOC exists in the initial state.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'compliance_manual'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set up default styles for the document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title Page ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('CORPORATE COMPLIANCE MANUAL')
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('Meridian Financial Services, Inc.')
    subtitle_run.font.size = Pt(16)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_run = version_para.add_run('Version 4.2 | Effective Date: January 1, 2025')
    version_run.font.size = Pt(12)
    version_run.italic = True

    doc.add_paragraph()  # spacer

    dept_para = doc.add_paragraph()
    dept_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept_run = dept_para.add_run('Compliance and Risk Management Department')
    dept_run.font.size = Pt(12)

    doc.add_page_break()

    # --- HEADING 1: Introduction (H1 #1) ---
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'This Corporate Compliance Manual ("Manual") has been prepared by Meridian Financial Services, '
        'Inc. to provide guidance on regulatory requirements, internal policies, and ethical standards '
        'that govern our business operations. All employees, contractors, and affiliated parties are '
        'required to familiarize themselves with this Manual and adhere to its provisions.'
    )

    # H2 #1
    doc.add_heading('Purpose and Scope', level=2)
    doc.add_paragraph(
        'The purpose of this Manual is to establish a comprehensive compliance framework that ensures '
        'adherence to applicable laws, regulations, and company policies. It applies to all business '
        'units, subsidiaries, and joint ventures in which Meridian Financial Services holds a '
        'controlling interest or management responsibility.'
    )
    doc.add_paragraph(
        'The scope encompasses domestic and international operations, covering regulatory compliance, '
        'anti-money laundering (AML) procedures, data protection requirements, and ethical conduct standards.'
    )

    # H2 #2
    doc.add_heading('Regulatory Framework Overview', level=2)
    doc.add_paragraph(
        'Meridian Financial Services operates under the oversight of multiple regulatory bodies, '
        'including the Securities and Exchange Commission (SEC), the Financial Industry Regulatory '
        'Authority (FINRA), and relevant state banking regulators. Our compliance program is designed '
        'to meet or exceed the requirements established by each of these authorities.'
    )

    # H3 #1
    doc.add_heading('Key Regulatory Authorities', level=3)
    doc.add_paragraph(
        'The primary regulatory authorities with jurisdiction over our operations include the SEC, '
        'FINRA, the Consumer Financial Protection Bureau (CFPB), the Office of the Comptroller of '
        'the Currency (OCC), and the Federal Reserve Board. Each authority has specific reporting, '
        'examination, and disclosure requirements that must be satisfied on an ongoing basis.'
    )

    # H2 #3
    doc.add_heading('Compliance Program Structure', level=2)
    doc.add_paragraph(
        'Our compliance program is structured around the three lines of defense model: business line '
        'management as the first line, the compliance function as the second line, and internal audit '
        'as the third line. This structure ensures that compliance responsibilities are clearly '
        'distributed and that independent oversight is maintained.'
    )

    # H3 #2
    doc.add_heading('Chief Compliance Officer Responsibilities', level=3)
    doc.add_paragraph(
        'The Chief Compliance Officer (CCO) is responsible for the overall design, implementation, '
        'and oversight of the compliance program. The CCO reports directly to the Board of Directors '
        'and maintains independence from business line management to ensure objective oversight and '
        'unbiased reporting.'
    )

    doc.add_page_break()

    # --- HEADING 1: Code of Business Conduct (H1 #2) ---
    doc.add_heading('Code of Business Conduct', level=1)
    doc.add_paragraph(
        'The Code of Business Conduct establishes the ethical standards and behavioral expectations '
        'for all employees of Meridian Financial Services. Adherence to this Code is a condition of '
        'employment and violations may result in disciplinary action up to and including termination.'
    )

    # H2 #4
    doc.add_heading('Ethical Standards and Integrity', level=2)
    doc.add_paragraph(
        'All employees are expected to conduct themselves with the highest standards of integrity, '
        'honesty, and professionalism. This includes accurate reporting of financial information, '
        'transparent communication with clients and regulators, and avoidance of deceptive or '
        'misleading practices in any form.'
    )

    # H3 #3
    doc.add_heading('Conflicts of Interest Policy', level=3)
    doc.add_paragraph(
        'Employees must avoid situations where personal interests conflict with those of the company '
        'or its clients. This includes outside employment, investments in competitors, and personal '
        'relationships with vendors or clients. All potential conflicts must be disclosed promptly '
        'to the compliance department using the Conflict Disclosure Form CF-100.'
    )

    # H3 #4
    doc.add_heading('Gift and Entertainment Guidelines', level=3)
    doc.add_paragraph(
        'Employees may not accept or offer gifts, entertainment, or other benefits that could '
        'influence business decisions or create the appearance of impropriety. Gifts with a value '
        'exceeding $75 must be reported to the compliance department. Entertainment expenses must '
        'be pre-approved for values exceeding $150 per person.'
    )

    # H2 #5
    doc.add_heading('Insider Trading Prevention', level=2)
    doc.add_paragraph(
        'Meridian Financial Services maintains a strict policy prohibiting insider trading and the '
        'misuse of material non-public information (MNPI). All employees who may have access to MNPI '
        'are subject to enhanced trading restrictions and must obtain pre-clearance before executing '
        'any personal securities transactions.'
    )

    # H3 #5
    doc.add_heading('Information Barriers', level=3)
    doc.add_paragraph(
        'Physical and logical information barriers ("Chinese walls") are maintained between business '
        'units that may have access to MNPI and those involved in public market activities. These '
        'barriers prevent the unauthorized flow of confidential information and are reviewed annually '
        'for effectiveness.'
    )

    doc.add_page_break()

    # --- HEADING 1: Anti-Money Laundering (H1 #3) ---
    doc.add_heading('Anti-Money Laundering Program', level=1)
    doc.add_paragraph(
        'The Anti-Money Laundering (AML) program is designed to detect, prevent, and report '
        'potential money laundering, terrorist financing, and other financial crimes. The program '
        'complies with the Bank Secrecy Act (BSA), the USA PATRIOT Act, and FinCEN guidelines.'
    )

    # H2 #6
    doc.add_heading('Customer Due Diligence', level=2)
    doc.add_paragraph(
        'Customer Due Diligence (CDD) procedures are applied to all new clients and reviewed '
        'periodically for existing clients. CDD includes identity verification, beneficial ownership '
        'determination, and risk assessment. Enhanced Due Diligence (EDD) is required for '
        'high-risk customers, politically exposed persons, and correspondent banking relationships.'
    )

    # H3 #6
    doc.add_heading('Know Your Customer Requirements', level=3)
    doc.add_paragraph(
        'KYC requirements mandate the collection and verification of customer identification '
        'information at account opening. Required information includes full legal name, date of birth, '
        'residential address, government-issued identification number, and source of funds for '
        'accounts exceeding specified thresholds. Documentation must be retained for a minimum of '
        'five years following account closure.'
    )

    # H3 #7
    doc.add_heading('Beneficial Ownership Procedures', level=3)
    doc.add_paragraph(
        'For legal entity customers, Meridian Financial Services must identify and verify all '
        'individuals who own 25% or more of the entity and any individual with significant managerial '
        'control. Beneficial ownership information is collected using the standard Certification Form '
        'BO-200 and verified against government databases and third-party sources.'
    )

    # H2 #7
    doc.add_heading('Transaction Monitoring', level=2)
    doc.add_paragraph(
        'An automated transaction monitoring system is deployed to identify unusual or suspicious '
        'activity patterns. The system applies scenario-based rules and behavioral analytics to flag '
        'transactions for review. All alerts are reviewed by trained AML analysts within 48 hours '
        'of generation.'
    )

    # H3 #8
    doc.add_heading('Suspicious Activity Reporting', level=3)
    doc.add_paragraph(
        'When suspicious activity is identified, a Suspicious Activity Report (SAR) must be filed '
        'with FinCEN within 30 calendar days of the date on which suspicious activity was initially '
        'detected. SAR filings are strictly confidential and may not be disclosed to the subject of '
        'the report. Employees who make good-faith SAR filings are protected from civil liability.'
    )

    doc.add_page_break()

    # --- HEADING 1: Data Protection and Privacy (H1 #4) ---
    doc.add_heading('Data Protection and Privacy', level=1)
    doc.add_paragraph(
        'Meridian Financial Services is committed to protecting the privacy and security of customer '
        'and employee data in accordance with applicable data protection laws, including the '
        'Gramm-Leach-Bliley Act (GLBA), the California Consumer Privacy Act (CCPA), and the '
        'General Data Protection Regulation (GDPR) for EU-related operations.'
    )

    # H2 #8
    doc.add_heading('Data Classification and Handling', level=2)
    doc.add_paragraph(
        'All data handled by Meridian Financial Services is classified into four categories: '
        'Public, Internal, Confidential, and Restricted. Each classification level carries specific '
        'handling, storage, transmission, and disposal requirements that must be followed by all '
        'employees who access data in that category.'
    )

    # H3 #9
    doc.add_heading('Personal Data Processing Standards', level=3)
    doc.add_paragraph(
        'Personal data may only be collected and processed for specified, legitimate purposes with '
        'appropriate legal basis. Data minimization principles require that only data necessary for '
        'the stated purpose be collected. Retention periods are defined in the Data Retention Schedule '
        'DR-500 and must be strictly observed.'
    )

    # H2 #9
    doc.add_heading('Cybersecurity Requirements', level=2)
    doc.add_paragraph(
        'The cybersecurity program establishes technical and organizational measures to protect '
        'information systems and sensitive data from unauthorized access, disclosure, or destruction. '
        'Requirements include multi-factor authentication for all system access, encryption of data '
        'at rest and in transit, and regular penetration testing.'
    )

    # H3 #10
    doc.add_heading('Incident Response Procedures', level=3)
    doc.add_paragraph(
        'In the event of a data breach or security incident, the Incident Response Team (IRT) must '
        'be activated within two hours of detection. Regulatory notification requirements vary by '
        'jurisdiction but generally require reporting within 72 hours. Affected customers must be '
        'notified promptly in accordance with applicable breach notification laws.'
    )

    doc.add_page_break()

    # --- HEADING 1: Securities Compliance (H1 #5) ---
    doc.add_heading('Securities Compliance', level=1)
    doc.add_paragraph(
        'The securities compliance program addresses requirements arising from the Securities Act of '
        '1933, the Securities Exchange Act of 1934, the Investment Advisers Act of 1940, and related '
        'regulations. This section applies primarily to registered representatives, investment '
        'advisers, and employees involved in securities activities.'
    )

    # H2 #10
    doc.add_heading('Suitability and Best Interest Standards', level=2)
    doc.add_paragraph(
        'All recommendations made to retail customers must satisfy Regulation Best Interest (Reg BI) '
        'standards, which require that recommendations be in the best interest of the customer at the '
        'time of recommendation. Investment advisers are subject to the fiduciary duty standard, '
        'which imposes a higher obligation of loyalty and care.'
    )

    # H3 #11
    doc.add_heading('Customer Profile Requirements', level=3)
    doc.add_paragraph(
        'Prior to making any recommendation, a complete customer profile must be obtained and '
        'documented. The profile must include investment objectives, risk tolerance, time horizon, '
        'financial situation, tax status, and investment experience. Profiles must be reviewed and '
        'updated at least annually or when significant changes occur.'
    )

    # H2 #11
    doc.add_heading('Trade Surveillance and Monitoring', level=2)
    doc.add_paragraph(
        'Automated trade surveillance monitors all securities transactions for potential violations, '
        'including excessive trading, front-running, market manipulation, and unauthorized trading. '
        'Exception reports are reviewed daily by compliance staff and escalated as appropriate.'
    )

    # H3 #12
    doc.add_heading('Order Handling and Best Execution', level=3)
    doc.add_paragraph(
        'Orders must be handled in accordance with applicable regulations regarding order handling, '
        'best execution, and customer order protection. Order routing decisions must prioritize '
        'customer interests, and execution quality is reviewed quarterly to ensure compliance with '
        'best execution obligations.'
    )

    doc.add_page_break()

    # --- HEADING 1: Human Resources Compliance (H1 #6) ---
    doc.add_heading('Human Resources Compliance', level=1)
    doc.add_paragraph(
        'Human resources compliance addresses employment-related regulatory requirements and '
        'internal policies governing the employer-employee relationship. This section covers '
        'equal employment opportunity, workplace safety, compensation practices, and employee '
        'training requirements.'
    )

    # H2 #12
    doc.add_heading('Equal Employment and Non-Discrimination', level=2)
    doc.add_paragraph(
        'Meridian Financial Services is an equal opportunity employer committed to maintaining a '
        'workplace free of discrimination and harassment. Employment decisions are based solely on '
        'legitimate business factors, and all employees are entitled to a respectful and inclusive '
        'work environment regardless of protected characteristics.'
    )

    # H3 #13
    doc.add_heading('Anti-Harassment Policy', level=3)
    doc.add_paragraph(
        'Workplace harassment, including sexual harassment, is strictly prohibited. All employees '
        'are required to complete anti-harassment training annually. Reports of harassment must be '
        'investigated promptly and confidentially. Retaliation against employees who report '
        'harassment in good faith is prohibited and constitutes a separate policy violation.'
    )

    # H2 #13
    doc.add_heading('Employee Training and Certification', level=2)
    doc.add_paragraph(
        'Mandatory training programs are required for all employees on an annual basis, covering '
        'compliance topics including AML, information security, code of conduct, and role-specific '
        'regulatory requirements. Training completion is tracked through the Learning Management '
        'System (LMS) and non-completion may result in disciplinary action.'
    )

    # H3 #14
    doc.add_heading('Licensing and Registration Requirements', level=3)
    doc.add_paragraph(
        'Employees engaged in regulated activities must maintain current licenses and registrations '
        'as required by applicable laws and regulations. This includes Series 7, Series 63, Series 65, '
        'and other FINRA licenses as applicable to job function. The compliance department maintains '
        'a license tracking system to ensure timely renewal and continuing education completion.'
    )

    doc.add_page_break()

    # --- HEADING 1: Vendor and Third-Party Management (H1 #7) ---
    doc.add_heading('Vendor and Third-Party Management', level=1)
    doc.add_paragraph(
        'The vendor management program governs relationships with external service providers, '
        'including technology vendors, outsourcing partners, and professional service firms. '
        'Third-party risk management is an integral component of the overall compliance and risk '
        'management framework.'
    )

    # H2 #14
    doc.add_heading('Vendor Due Diligence', level=2)
    doc.add_paragraph(
        'All vendors accessing company systems, data, or facilities must complete a due diligence '
        'review prior to engagement. Due diligence scope is determined by vendor risk tier and '
        'includes financial stability assessment, regulatory compliance review, information security '
        'questionnaire, and reference checks.'
    )

    # H2 #15
    doc.add_heading('Contract Requirements and Service Standards', level=2)
    doc.add_paragraph(
        'All vendor contracts must include standard compliance provisions, including data protection '
        'requirements, audit rights, regulatory cooperation obligations, and termination provisions. '
        'Service Level Agreements (SLAs) must specify performance standards, remediation procedures, '
        'and escalation protocols for service failures.'
    )

    # H3 #15
    doc.add_heading('Subcontractor Oversight', level=3)
    doc.add_paragraph(
        'Vendors who engage subcontractors to perform services on behalf of Meridian Financial '
        'Services must notify the company and obtain approval. Subcontractors are subject to the '
        'same due diligence and contractual requirements as primary vendors. Flow-down provisions '
        'in primary vendor contracts must ensure that compliance obligations are passed through '
        'to all levels of the supply chain.'
    )

    doc.add_page_break()

    # --- HEADING 1: Regulatory Reporting and Recordkeeping (H1 #8) ---
    doc.add_heading('Regulatory Reporting and Recordkeeping', level=1)
    doc.add_paragraph(
        'Accurate and timely regulatory reporting is a fundamental compliance obligation. This '
        'section covers the primary reporting requirements applicable to Meridian Financial Services '
        'operations, record retention standards, and examination readiness procedures.'
    )

    # H2 #16
    doc.add_heading('Financial Reporting Requirements', level=2)
    doc.add_paragraph(
        'Financial reports must be prepared in accordance with Generally Accepted Accounting '
        'Principles (GAAP) and applicable SEC reporting requirements. All public filings, including '
        '10-K annual reports, 10-Q quarterly reports, and 8-K current reports, must be reviewed '
        'by the legal and compliance departments prior to submission.'
    )

    # H2 #17
    doc.add_heading('Record Retention Standards', level=2)
    doc.add_paragraph(
        'Business records must be retained in accordance with the Record Retention Schedule '
        'published by the compliance department. Retention periods range from three years for '
        'routine correspondence to seven years for transaction records and indefinitely for '
        'corporate formation documents. Electronic records must be maintained in approved systems '
        'with appropriate access controls and backup procedures.'
    )

    # H2 #18
    doc.add_heading('Examination and Audit Management', level=2)
    doc.add_paragraph(
        'Regulatory examinations must be managed in close coordination with the compliance and '
        'legal departments. All communications with examiners must be approved by the CCO or '
        'General Counsel. Document requests must be fulfilled accurately and promptly, and all '
        'responses must be reviewed for completeness and accuracy before submission.'
    )

    # H2 #19
    doc.add_heading('Whistleblower and Reporting Mechanisms', level=2)
    doc.add_paragraph(
        'Employees are encouraged to report suspected compliance violations through appropriate '
        'channels, including the anonymous Ethics Hotline (1-800-MERIDIAN), direct reporting to '
        'the compliance department, or reporting to the Board Audit Committee. Reports may be made '
        'anonymously and are investigated confidentially. Retaliation against whistleblowers is '
        'strictly prohibited.'
    )

    # H2 #20
    doc.add_heading('Compliance Program Testing and Assessment', level=2)
    doc.add_paragraph(
        'The compliance program is subject to regular testing and assessment to evaluate '
        'effectiveness and identify areas for improvement. Annual compliance risk assessments '
        'prioritize resources and identify emerging risks. Monthly compliance metrics are reported '
        'to senior management and quarterly summaries are presented to the Board of Directors.'
    )

    # Appendix section to pad to ~20 pages
    doc.add_page_break()

    doc.add_paragraph('APPENDIX A: COMPLIANCE FORMS AND TEMPLATES').runs[0].bold = True
    doc.add_paragraph(
        'This appendix contains standardized forms referenced throughout this Manual, including '
        'the Conflict Disclosure Form CF-100, the Beneficial Ownership Certification Form BO-200, '
        'and the Trade Pre-Clearance Form TC-300. Forms are available in electronic format through '
        'the compliance intranet portal. Completed forms must be submitted to the compliance '
        'department within the timeframes specified in the applicable policy sections.'
    )

    doc.add_page_break()

    doc.add_paragraph('APPENDIX B: REGULATORY REFERENCE GUIDE').runs[0].bold = True
    doc.add_paragraph(
        'This appendix provides a summary of key regulations applicable to Meridian Financial '
        'Services operations, including citation references, effective dates, and brief descriptions '
        'of primary requirements. This guide is intended as a quick reference tool and should not '
        'be relied upon as legal advice. Consult the legal department for specific regulatory '
        'interpretation questions.'
    )

    doc.add_page_break()

    doc.add_paragraph('APPENDIX C: TRAINING REQUIREMENTS MATRIX').runs[0].bold = True
    doc.add_paragraph(
        'The Training Requirements Matrix provides a comprehensive listing of all mandatory and '
        'recommended training programs, including frequency, applicable employee groups, and '
        'completion deadlines. Training completion is tracked automatically through the LMS system '
        'and non-completion triggers automated escalation notices.'
    )

    doc.add_page_break()

    doc.add_paragraph('DOCUMENT CONTROL').runs[0].bold = True
    doc.add_paragraph(
        'Document Title: Corporate Compliance Manual\n'
        'Version: 4.2\n'
        'Effective Date: January 1, 2025\n'
        'Review Cycle: Annual\n'
        'Owner: Chief Compliance Officer\n'
        'Approver: Board of Directors\n'
        'Classification: Internal Use\n'
        'Next Review Date: January 1, 2026'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
