"""
Initial Setup: Contract document with 8 pages, ending with witness clause, no signature block.
Task ID: writer_pd_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title Page (Page 1) ---
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_heading("PROFESSIONAL SERVICES AGREEMENT", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")

    subtitle = doc.add_paragraph("Between")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph("")

    party_a_name = doc.add_paragraph("Meridian Technology Solutions, Inc.")
    party_a_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_a = party_a_name.runs[0]
    run_a.bold = True
    run_a.font.size = Pt(16)

    doc.add_paragraph("")

    and_text = doc.add_paragraph("and")
    and_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    and_text.runs[0].font.size = Pt(14)

    doc.add_paragraph("")

    party_b_name = doc.add_paragraph("Cascade Digital Ventures, LLC")
    party_b_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_b = party_b_name.runs[0]
    run_b.bold = True
    run_b.font.size = Pt(16)

    for _ in range(4):
        doc.add_paragraph("")

    date_line = doc.add_paragraph("Effective Date: March 15, 2025")
    date_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_line.runs[0].font.size = Pt(12)

    ref_line = doc.add_paragraph("Contract Reference: PSA-2025-0472")
    ref_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ref_line.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # --- Page 2: Table of Contents / Recitals ---
    doc.add_heading("TABLE OF CONTENTS", level=1)
    toc_items = [
        "Article I — Definitions",
        "Article II — Scope of Services",
        "Article III — Compensation and Payment Terms",
        "Article IV — Term and Termination",
        "Article V — Intellectual Property Rights",
        "Article VI — Confidentiality",
        "Article VII — Representations and Warranties",
        "Article VIII — Indemnification",
        "Article IX — Limitation of Liability",
        "Article X — General Provisions",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph("")
    doc.add_heading("RECITALS", level=1)

    doc.add_paragraph(
        "WHEREAS, Meridian Technology Solutions, Inc. (hereinafter referred to as "
        "\"Party A\") is a corporation duly organized under the laws of the State of "
        "Delaware, with its principal place of business at 4500 Innovation Drive, "
        "Suite 300, San Jose, California 95134;"
    )

    doc.add_paragraph(
        "WHEREAS, Cascade Digital Ventures, LLC (hereinafter referred to as "
        "\"Party B\") is a limited liability company organized under the laws of the "
        "State of Washington, with its principal place of business at 1200 Waterfront "
        "Boulevard, Suite 800, Seattle, Washington 98101;"
    )

    doc.add_paragraph(
        "WHEREAS, Party A desires to engage Party B to provide certain professional "
        "consulting and software development services, and Party B desires to provide "
        "such services to Party A, subject to the terms and conditions set forth herein;"
    )

    doc.add_paragraph(
        "NOW, THEREFORE, in consideration of the mutual covenants, promises, and "
        "agreements contained herein, and for other good and valuable consideration, "
        "the receipt and sufficiency of which are hereby acknowledged, the parties "
        "agree as follows:"
    )

    doc.add_page_break()

    # --- Page 3: Article I - Definitions ---
    doc.add_heading("ARTICLE I — DEFINITIONS", level=1)

    definitions = [
        ('"Confidential Information"', 'means any non-public, proprietary, or trade secret information disclosed by either party to the other, whether orally, in writing, or by inspection of tangible objects, including but not limited to business plans, financial data, customer lists, technical specifications, source code, algorithms, and marketing strategies.'),
        ('"Deliverables"', 'means the work product, including software code, documentation, reports, designs, and other materials that Party B is required to deliver to Party A pursuant to each Statement of Work.'),
        ('"Effective Date"', 'means March 15, 2025, the date on which this Agreement becomes binding upon both parties.'),
        ('"Intellectual Property"', 'means all patents, copyrights, trademarks, trade secrets, and other proprietary rights in any jurisdiction, including all applications and registrations relating thereto.'),
        ('"Services"', 'means the professional consulting, software development, system integration, and related technical services to be provided by Party B as described in the applicable Statement of Work.'),
        ('"Statement of Work" or "SOW"', 'means a written document executed by both parties that describes the specific Services to be performed, Deliverables, timeline, milestones, and fees applicable to a particular project or engagement.'),
    ]

    for term, definition in definitions:
        p = doc.add_paragraph()
        run_term = p.add_run(term)
        run_term.bold = True
        p.add_run(f" {definition}")
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # --- Page 4: Article II & III ---
    doc.add_heading("ARTICLE II — SCOPE OF SERVICES", level=1)

    doc.add_paragraph(
        "2.1 Party B shall provide the Services described in each Statement of Work "
        "executed by both parties. Each SOW shall be incorporated into and made a part "
        "of this Agreement. In the event of a conflict between the terms of this "
        "Agreement and any SOW, the terms of this Agreement shall prevail unless the "
        "SOW explicitly states otherwise."
    )
    doc.add_paragraph(
        "2.2 Party B shall perform the Services in a professional and workmanlike "
        "manner, consistent with industry standards and best practices. Party B shall "
        "assign qualified personnel with appropriate skills and experience to perform "
        "the Services."
    )
    doc.add_paragraph(
        "2.3 Party B shall comply with all applicable federal, state, and local laws, "
        "regulations, and ordinances in the performance of the Services. Party B shall "
        "obtain and maintain all necessary licenses, permits, and certifications required "
        "to perform the Services."
    )
    doc.add_paragraph(
        "2.4 Any changes to the scope of Services described in a SOW shall be documented "
        "in a written change order signed by authorized representatives of both parties. "
        "No additional Services shall be performed, and no additional fees shall be "
        "incurred, without prior written approval from Party A."
    )

    doc.add_paragraph("")

    doc.add_heading("ARTICLE III — COMPENSATION AND PAYMENT TERMS", level=1)

    doc.add_paragraph(
        "3.1 Party A shall compensate Party B for the Services at the rates and in the "
        "amounts specified in each applicable SOW. Unless otherwise specified in a SOW, "
        "Party B shall invoice Party A monthly in arrears for Services performed during "
        "the preceding calendar month."
    )
    doc.add_paragraph(
        "3.2 All invoices shall be payable within thirty (30) days of receipt. Late "
        "payments shall bear interest at the rate of one and one-half percent (1.5%) "
        "per month or the maximum rate permitted by applicable law, whichever is less."
    )
    doc.add_paragraph(
        "3.3 Party A shall reimburse Party B for reasonable, pre-approved travel and "
        "out-of-pocket expenses incurred in connection with the performance of Services, "
        "provided that Party B submits itemized receipts and documentation within fifteen "
        "(15) days of incurring such expenses."
    )

    doc.add_page_break()

    # --- Page 5: Article IV & V ---
    doc.add_heading("ARTICLE IV — TERM AND TERMINATION", level=1)

    doc.add_paragraph(
        "4.1 This Agreement shall commence on the Effective Date and shall continue for "
        "an initial term of twenty-four (24) months, unless earlier terminated in "
        "accordance with the provisions of this Article. Upon expiration of the initial "
        "term, this Agreement shall automatically renew for successive twelve (12) month "
        "periods unless either party provides written notice of non-renewal at least "
        "ninety (90) days prior to the end of the then-current term."
    )
    doc.add_paragraph(
        "4.2 Either party may terminate this Agreement for cause upon sixty (60) days' "
        "written notice to the other party if the other party materially breaches any "
        "provision of this Agreement and fails to cure such breach within the sixty (60) "
        "day notice period."
    )
    doc.add_paragraph(
        "4.3 Either party may terminate this Agreement without cause upon one hundred "
        "twenty (120) days' prior written notice to the other party. In the event of "
        "termination without cause, Party A shall pay Party B for all Services performed "
        "and expenses incurred through the effective date of termination."
    )
    doc.add_paragraph(
        "4.4 Upon termination or expiration of this Agreement, Party B shall promptly "
        "deliver to Party A all Deliverables completed or in progress, all Confidential "
        "Information of Party A, and all other materials belonging to Party A."
    )

    doc.add_paragraph("")

    doc.add_heading("ARTICLE V — INTELLECTUAL PROPERTY RIGHTS", level=1)

    doc.add_paragraph(
        "5.1 All Deliverables created by Party B in the performance of the Services "
        "shall be considered works made for hire and shall be the sole and exclusive "
        "property of Party A. To the extent any Deliverable does not qualify as a work "
        "made for hire, Party B hereby irrevocably assigns to Party A all right, title, "
        "and interest in and to such Deliverable."
    )
    doc.add_paragraph(
        "5.2 Party B retains all rights in its pre-existing intellectual property, "
        "tools, methodologies, and know-how that existed prior to the Effective Date or "
        "that are developed independently of the Services. Party B hereby grants Party A "
        "a non-exclusive, perpetual, royalty-free license to use any such pre-existing "
        "intellectual property incorporated into the Deliverables."
    )

    doc.add_page_break()

    # --- Page 6: Article VI & VII ---
    doc.add_heading("ARTICLE VI — CONFIDENTIALITY", level=1)

    doc.add_paragraph(
        "6.1 Each party agrees to hold in strict confidence all Confidential Information "
        "received from the other party and to use such information solely for the purposes "
        "contemplated by this Agreement. Each party shall protect the other party's "
        "Confidential Information using the same degree of care it uses to protect its own "
        "confidential information, but in no event less than reasonable care."
    )
    doc.add_paragraph(
        "6.2 The obligations of confidentiality shall not apply to information that: "
        "(a) is or becomes publicly available through no fault of the receiving party; "
        "(b) was known to the receiving party prior to disclosure; (c) is independently "
        "developed by the receiving party without use of or reference to the disclosing "
        "party's Confidential Information; or (d) is rightfully obtained from a third "
        "party without restriction on disclosure."
    )
    doc.add_paragraph(
        "6.3 The obligations of confidentiality set forth in this Article shall survive "
        "the termination or expiration of this Agreement for a period of five (5) years."
    )

    doc.add_paragraph("")

    doc.add_heading("ARTICLE VII — REPRESENTATIONS AND WARRANTIES", level=1)

    doc.add_paragraph(
        "7.1 Each party represents and warrants that: (a) it has the full power and "
        "authority to enter into this Agreement and to perform its obligations hereunder; "
        "(b) the execution and performance of this Agreement does not and will not conflict "
        "with any other agreement to which it is a party; and (c) it shall comply with all "
        "applicable laws in the performance of its obligations."
    )
    doc.add_paragraph(
        "7.2 Party B further represents and warrants that: (a) the Services will be "
        "performed in a professional manner consistent with industry standards; "
        "(b) the Deliverables will conform to the specifications set forth in the "
        "applicable SOW; (c) the Deliverables will not infringe upon any third party's "
        "intellectual property rights; and (d) Party B has obtained all necessary consents "
        "and licenses to perform the Services."
    )

    doc.add_page_break()

    # --- Page 7: Article VIII, IX ---
    doc.add_heading("ARTICLE VIII — INDEMNIFICATION", level=1)

    doc.add_paragraph(
        "8.1 Party B shall indemnify, defend, and hold harmless Party A and its officers, "
        "directors, employees, and agents from and against any and all claims, damages, "
        "losses, liabilities, costs, and expenses (including reasonable attorneys' fees) "
        "arising out of or relating to: (a) any breach by Party B of its representations, "
        "warranties, or obligations under this Agreement; (b) any negligent or willful "
        "misconduct by Party B or its personnel in the performance of the Services; or "
        "(c) any claim that the Deliverables infringe upon the intellectual property rights "
        "of any third party."
    )
    doc.add_paragraph(
        "8.2 Party A shall indemnify, defend, and hold harmless Party B and its officers, "
        "directors, employees, and agents from and against any and all claims, damages, "
        "losses, liabilities, costs, and expenses (including reasonable attorneys' fees) "
        "arising out of or relating to any breach by Party A of its representations, "
        "warranties, or obligations under this Agreement."
    )

    doc.add_paragraph("")

    doc.add_heading("ARTICLE IX — LIMITATION OF LIABILITY", level=1)

    doc.add_paragraph(
        "9.1 EXCEPT FOR OBLIGATIONS UNDER ARTICLE VI (CONFIDENTIALITY) AND ARTICLE VIII "
        "(INDEMNIFICATION), NEITHER PARTY SHALL BE LIABLE TO THE OTHER FOR ANY INDIRECT, "
        "INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING BUT NOT LIMITED "
        "TO LOSS OF PROFITS, REVENUE, DATA, OR BUSINESS OPPORTUNITIES, REGARDLESS OF THE "
        "CAUSE OF ACTION OR THEORY OF LIABILITY, EVEN IF SUCH PARTY HAS BEEN ADVISED OF "
        "THE POSSIBILITY OF SUCH DAMAGES."
    )
    doc.add_paragraph(
        "9.2 THE TOTAL AGGREGATE LIABILITY OF EITHER PARTY UNDER THIS AGREEMENT SHALL NOT "
        "EXCEED THE TOTAL AMOUNT OF FEES PAID OR PAYABLE BY PARTY A TO PARTY B DURING THE "
        "TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM."
    )

    doc.add_page_break()

    # --- Page 8: Article X and Witness Clause ---
    doc.add_heading("ARTICLE X — GENERAL PROVISIONS", level=1)

    doc.add_paragraph(
        "10.1 Governing Law. This Agreement shall be governed by and construed in "
        "accordance with the laws of the State of California, without regard to its "
        "conflict of laws provisions. Any dispute arising under or relating to this "
        "Agreement shall be resolved exclusively in the state or federal courts located "
        "in Santa Clara County, California."
    )
    doc.add_paragraph(
        "10.2 Entire Agreement. This Agreement, together with all SOWs and exhibits "
        "attached hereto, constitutes the entire agreement between the parties with "
        "respect to the subject matter hereof and supersedes all prior and contemporaneous "
        "agreements, negotiations, and understandings, whether written or oral."
    )
    doc.add_paragraph(
        "10.3 Amendment. No amendment, modification, or waiver of any provision of this "
        "Agreement shall be effective unless made in writing and signed by authorized "
        "representatives of both parties."
    )
    doc.add_paragraph(
        "10.4 Assignment. Neither party may assign or transfer this Agreement or any "
        "rights or obligations hereunder without the prior written consent of the other "
        "party, except that either party may assign this Agreement in connection with a "
        "merger, acquisition, or sale of substantially all of its assets."
    )
    doc.add_paragraph(
        "10.5 Severability. If any provision of this Agreement is held to be invalid or "
        "unenforceable, the remaining provisions shall continue in full force and effect, "
        "and the invalid or unenforceable provision shall be modified to the minimum extent "
        "necessary to make it valid and enforceable."
    )
    doc.add_paragraph(
        "10.6 Notices. All notices required or permitted under this Agreement shall be in "
        "writing and shall be deemed given when delivered personally, sent by certified "
        "mail (return receipt requested), or sent by nationally recognized overnight "
        "courier to the addresses set forth above or to such other address as a party "
        "may designate in writing."
    )
    doc.add_paragraph(
        "10.7 Force Majeure. Neither party shall be liable for any failure or delay in "
        "performing its obligations under this Agreement to the extent that such failure "
        "or delay results from circumstances beyond the party's reasonable control, "
        "including but not limited to acts of God, natural disasters, war, terrorism, "
        "pandemics, government orders, or interruption of utility services."
    )
    doc.add_paragraph(
        "10.8 Counterparts. This Agreement may be executed in counterparts, each of which "
        "shall be deemed an original and all of which together shall constitute one and "
        "the same instrument. Electronic signatures shall be deemed valid and binding."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    # The witness clause - final text before the (missing) signature block
    witness = doc.add_paragraph(
        "IN WITNESS WHEREOF, the parties have executed this agreement."
    )
    witness.runs[0].bold = True
    witness.runs[0].font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
