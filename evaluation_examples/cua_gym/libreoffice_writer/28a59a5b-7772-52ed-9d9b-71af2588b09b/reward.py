"""
Reward Script: Insert signature block with two-column layout for Party A and Party B
Task ID: writer_pd_005
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Signature table exists with 2-column layout
  Component 2 (0.25): Party A header in left column, Party B header in right column
  Component 3 (0.30): All four fields (Signature, Printed Name, Title, Date) in both columns
  Component 4 (0.20): Signature lines present with adequate width
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_005'


def persist_app_state(domain: str):
    """Save any unsaved changes in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def find_signature_table(doc):
    """
    Find a table that serves as a signature block.
    A signature table must have at least 2 columns and contain
    references to 'Party A' and 'Party B' (or similar).
    """
    for ti, table in enumerate(doc.tables):
        if len(table.columns) < 2:
            continue
        full_text = ""
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text.lower() + " "
        if "party a" in full_text and "party b" in full_text:
            return ti, table
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: the witness text should still be present
    witness_found = any("in witness whereof" in p.text.lower() for p in doc.paragraphs)
    if not witness_found:
        print("PRECONDITION FAIL: 'IN WITNESS WHEREOF' text not found in document")
        print("REWARD: 0.0")
        return 0.0

    # Find the signature table
    table_idx, sig_table = find_signature_table(doc)

    # Component 1: Signature table exists with 2-column layout (0.25 points)
    try:
        if sig_table is not None and len(sig_table.columns) >= 2:
            print(f"PASS: Component 1 -- Signature table found (table index {table_idx}) "
                  f"with {len(sig_table.columns)} columns, {len(sig_table.rows)} rows (0.25 pts)")
            total_score += 0.25
        else:
            if sig_table is None:
                print("FAIL: Component 1 -- No signature table found with Party A/Party B references")
            else:
                print(f"FAIL: Component 1 -- Table found but has {len(sig_table.columns)} columns, need >= 2")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # If no table found, remaining checks cannot pass
    if sig_table is None:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Collect text from left column (col 0) and right column (col 1)
    left_texts = []
    right_texts = []
    for row in sig_table.rows:
        cells = row.cells
        if len(cells) >= 2:
            left_texts.append(cells[0].text.strip().lower())
            right_texts.append(cells[1].text.strip().lower())

    left_all = " ".join(left_texts)
    right_all = " ".join(right_texts)

    # Component 2: Party A in left column, Party B in right column (0.25 points)
    try:
        party_a_left = "party a" in left_all
        party_b_right = "party b" in right_all
        if party_a_left and party_b_right:
            print(f"PASS: Component 2 -- Party A found in left column, Party B in right column (0.25 pts)")
            total_score += 0.25
        else:
            details = []
            if not party_a_left:
                details.append("Party A not found in left column")
            if not party_b_right:
                details.append("Party B not found in right column")
            print(f"FAIL: Component 2 -- {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: All four fields present in both columns (0.30 points)
    # Fields: Signature, Printed Name, Title, Date
    # Each field found in both columns = 0.075 points (4 fields * 0.075 = 0.30)
    required_fields = ["signature", "printed name", "title", "date"]
    try:
        fields_score = 0.0
        for field in required_fields:
            in_left = field in left_all
            in_right = field in right_all
            if in_left and in_right:
                print(f"  PASS: Field '{field}' found in both columns")
                fields_score += 0.075
            else:
                missing = []
                if not in_left:
                    missing.append("left")
                if not in_right:
                    missing.append("right")
                print(f"  FAIL: Field '{field}' missing from: {', '.join(missing)}")

        if fields_score > 0:
            print(f"PASS: Component 3 -- Signature fields check ({fields_score:.3f} pts)")
            total_score += fields_score
        else:
            print("FAIL: Component 3 -- No required signature fields found in table")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Signature lines present with adequate width (0.20 points)
    # Check for underscore lines (___) in cells, need at least 15 underscores per line
    # (15 chars at ~0.35cm each approx = 5.25cm, meeting the 5cm requirement)
    try:
        lines_found = 0
        min_underscore_len = 15  # minimum underscores for a valid signature line
        for row in sig_table.rows:
            for cell in row.cells:
                text = cell.text
                # Find sequences of underscores
                underscore_matches = re.findall(r'_{%d,}' % min_underscore_len, text)
                if underscore_matches:
                    lines_found += 1

        # We expect at least 4 signature lines per column = 8 total
        # But be lenient: at least 4 total lines across both columns shows effort
        if lines_found >= 6:
            print(f"PASS: Component 4 -- {lines_found} signature lines found with adequate width (0.20 pts)")
            total_score += 0.20
        elif lines_found >= 3:
            partial = 0.10
            print(f"PARTIAL: Component 4 -- {lines_found} signature lines found (partial: {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 -- Only {lines_found} signature lines found, need at least 6")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
