"""
Initial Setup: Linguistics research paper with 4 body paragraphs and 2 bibliography entries
Task ID: osworld_writer_bibliography_crossref_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Exploring Language Structure: A Multi-Faceted Approach", level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author & affiliation ---
    author_para = doc.add_paragraph("Dr. Emily Harrington")
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in author_para.runs:
        run.font.italic = True
        run.font.size = Pt(12)

    affil_para = doc.add_paragraph("Department of Linguistics, University of Cambridge, Cambridge, UK")
    affil_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in affil_para.runs:
        run.font.size = Pt(10)

    doc.add_paragraph("")

    # --- Abstract ---
    abstract_heading = doc.add_heading("Abstract", level=1)
    abstract_text = doc.add_paragraph(
        "This paper investigates the intersection of syntax, pragmatics, and discourse analysis "
        "within modern linguistic theory. Drawing on contemporary frameworks, we examine how "
        "structural properties of language interact with communicative context and cognitive "
        "processes. Our analysis demonstrates that a unified approach incorporating both formal "
        "and functional perspectives yields richer explanatory power for cross-linguistic phenomena."
    )
    abstract_text.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")

    # --- Introduction heading ---
    doc.add_heading("1. Introduction", level=1)

    # --- Paragraph 1 (NO citation yet - agent will add (Wilson, 2019) after this) ---
    para1 = doc.add_paragraph(
        "The study of language structure has undergone significant transformation over the past "
        "two decades. Traditional generative approaches, which emphasized the formal properties "
        "of syntactic derivations, have been increasingly complemented by usage-based models "
        "that situate grammatical knowledge within broader patterns of language use. This "
        "theoretical convergence has opened productive avenues for examining how speakers "
        "acquire, process, and deploy linguistic structures in real-time communicative contexts. "
        "The implications for language pedagogy and computational modeling are profound, "
        "necessitating a re-examination of foundational assumptions about the architecture of "
        "the language faculty."
    )
    para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para1.paragraph_format.space_after = Pt(6)

    # --- Paragraph 2 (NO citation yet - agent will add (Zhang, 2020) after this) ---
    para2 = doc.add_paragraph(
        "Computational approaches to syntax have provided novel insights into the formal "
        "underpinnings of grammatical structure. By leveraging large-scale corpora and "
        "machine learning techniques, researchers have been able to identify statistical "
        "regularities that both confirm and challenge existing theoretical accounts. "
        "Dependency parsing algorithms, in particular, have shed light on the hierarchical "
        "relations governing sentence structure across typologically diverse languages. "
        "These findings raise important questions about the extent to which syntactic "
        "universals reflect cognitive constraints versus culturally-transmitted conventions, "
        "a debate that continues to animate theoretical discussions in the field."
    )
    para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para2.paragraph_format.space_after = Pt(6)

    # --- Paragraph 3 (NO citation yet - agent will add (Brown, 2021) after this) ---
    para3 = doc.add_paragraph(
        "Pragmatic theory has long sought to account for the gap between semantic content and "
        "communicative meaning. Speaker intent, contextual inference, and the negotiation of "
        "common ground all play crucial roles in determining how utterances are interpreted "
        "in natural discourse. The Gricean framework of conversational implicature, while "
        "foundational, has been supplemented by relevance-theoretic accounts that emphasize "
        "cognitive efficiency in pragmatic processing. Recent experimental work using "
        "eye-tracking and response-time measures has begun to clarify the neural substrates "
        "of pragmatic inference, providing an empirical basis for theoretical refinements "
        "that better capture the dynamic interplay between literal meaning and speaker intention."
    )
    para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para3.paragraph_format.space_after = Pt(6)

    # --- Paragraph 4 (NO (Taylor, 2022) citation yet - agent will insert it in this paragraph) ---
    para4 = doc.add_paragraph(
        "Discourse analysis offers a complementary perspective by focusing on the structural "
        "and functional organization of extended language use beyond the sentence level. "
        "Coherence relations, topic management, and rhetorical structure all contribute to "
        "the texture of well-formed discourse. Ethnographic approaches have enriched our "
        "understanding of how discourse practices vary across social and cultural settings, "
        "revealing that what counts as coherent or appropriate communication is deeply "
        "context-dependent. This recognition has significant implications for cross-cultural "
        "communication, language education, and the design of natural language processing "
        "systems intended to operate in diverse linguistic environments."
    )
    para4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para4.paragraph_format.space_after = Pt(6)

    # --- Conclusion heading ---
    doc.add_heading("2. Conclusion", level=1)
    conclusion = doc.add_paragraph(
        "This paper has surveyed key developments across syntax, computational linguistics, "
        "pragmatics, and discourse analysis, highlighting the productive tensions and "
        "complementarities among these subfields. Future research should continue to pursue "
        "integrative frameworks that honor both formal rigor and empirical breadth. "
        "Collaboration across theoretical traditions and methodological paradigms remains "
        "essential for advancing our understanding of the complex phenomenon of human language."
    )
    conclusion.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")

    # --- Bibliography section with 2 existing entries (NOT the 4 new ones) ---
    bib_heading = doc.add_heading("Bibliography", level=1)

    bib1 = doc.add_paragraph(
        "Chomsky, N. (1995). The Minimalist Program. MIT Press."
    )
    bib1.paragraph_format.left_indent = Pt(36)
    bib1.paragraph_format.first_line_indent = Pt(-36)

    bib2 = doc.add_paragraph(
        "Tomasello, M. (2003). Constructing a Language: A Usage-Based Theory of Language Acquisition. Harvard University Press."
    )
    bib2.paragraph_format.left_indent = Pt(36)
    bib2.paragraph_format.first_line_indent = Pt(-36)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
