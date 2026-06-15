"""
Reward Script: Add paragraph borders to RECITALS section
Task ID: writer_legal_042
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): First WHEREAS paragraph has a top border
  Component 2 (0.4): Last WHEREAS paragraph has a bottom border
  Component 3 (0.2): Border line properties are correct (single, ~0.5pt, dark color)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_042'


def get_border_info(para, side):
    """
    Extract border info for a given side ('top', 'bottom', 'left', 'right') from a paragraph.
    Returns dict with keys: val, sz, color, space  or None if no border found.
    """
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        return None
    border_el = pBdr.find(qn(f'w:{side}'))
    if border_el is None:
        return None
    return {
        'val': border_el.get(qn('w:val')),
        'sz': border_el.get(qn('w:sz')),
        'color': border_el.get(qn('w:color')),
        'space': border_el.get(qn('w:space')),
    }


def find_recitals_whereas_paras(doc):
    """
    Find the WHEREAS paragraphs in the RECITALS section.
    Returns list of (index, paragraph) tuples.
    """
    whereas_paras = []
    recitals_start = -1
    # First pass: find the RECITALS heading index
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if 'RECITAL' in text and (para.style.name.startswith('Heading') or text == 'RECITALS'):
            recitals_start = i
            break
    if recitals_start < 0:
        return whereas_paras
    # Second pass: collect WHEREAS paragraphs after RECITALS until next heading
    for i in range(recitals_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style.name.startswith('Heading'):
            break
        if para.text.strip().upper().startswith('WHEREAS'):
            whereas_paras.append((i, para))
    return whereas_paras


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the WHEREAS paragraphs in the RECITALS section
    whereas_paras = find_recitals_whereas_paras(doc)
    if len(whereas_paras) < 2:
        print(f"FAIL: Expected at least 2 WHEREAS paragraphs, found {len(whereas_paras)}")
        print("REWARD: 0.0")
        return 0.0

    first_idx, first_para = whereas_paras[0]
    last_idx, last_para = whereas_paras[-1]
    print(f"INFO: Found {len(whereas_paras)} WHEREAS paragraphs (indices: {[w[0] for w in whereas_paras]})")

    # Component 1: First WHEREAS paragraph has a top border (0.4 points)
    try:
        top_border = get_border_info(first_para, 'top')
        if top_border is not None and top_border['val'] is not None and top_border['val'] != 'none':
            print(f"PASS: Component 1 — First WHEREAS (para {first_idx}) has top border: val={top_border['val']}, sz={top_border['sz']}, color={top_border['color']} (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — First WHEREAS (para {first_idx}) has no top border (got: {top_border})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Last WHEREAS paragraph has a bottom border (0.4 points)
    try:
        bottom_border = get_border_info(last_para, 'bottom')
        if bottom_border is not None and bottom_border['val'] is not None and bottom_border['val'] != 'none':
            print(f"PASS: Component 2 — Last WHEREAS (para {last_idx}) has bottom border: val={bottom_border['val']}, sz={bottom_border['sz']}, color={bottom_border['color']} (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Last WHEREAS (para {last_idx}) has no bottom border (got: {bottom_border})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Border properties are correct — single line, approximately 0.5pt (sz<=8),
    # and dark color (0.2 points)
    # Only award if at least one border was found from components 1 or 2
    try:
        borders_correct = 0
        checks_done = 0

        if top_border is not None and top_border['val'] is not None and top_border['val'] != 'none':
            checks_done += 1
            val_ok = top_border['val'] == 'single'
            # sz is in 1/8 pt; 0.5pt = 4, allow up to 1pt = 8
            sz_ok = top_border['sz'] is not None and int(top_border['sz']) <= 8
            # color should be dark (000000 or auto or similar)
            color = (top_border['color'] or '').lower()
            color_ok = color in ('000000', 'auto', '000001', '') or color.startswith('00')
            if val_ok and sz_ok and color_ok:
                borders_correct += 1

        if bottom_border is not None and bottom_border['val'] is not None and bottom_border['val'] != 'none':
            checks_done += 1
            val_ok = bottom_border['val'] == 'single'
            sz_ok = bottom_border['sz'] is not None and int(bottom_border['sz']) <= 8
            color = (bottom_border['color'] or '').lower()
            color_ok = color in ('000000', 'auto', '000001', '') or color.startswith('00')
            if val_ok and sz_ok and color_ok:
                borders_correct += 1

        if checks_done > 0 and borders_correct == checks_done:
            print(f"PASS: Component 3 — Border properties correct: single line, thin, dark color (0.2 pts)")
            total_score += 0.2
        elif checks_done > 0:
            print(f"FAIL: Component 3 — Some border properties incorrect ({borders_correct}/{checks_done} correct)")
        else:
            print(f"FAIL: Component 3 — No borders found to verify properties")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
