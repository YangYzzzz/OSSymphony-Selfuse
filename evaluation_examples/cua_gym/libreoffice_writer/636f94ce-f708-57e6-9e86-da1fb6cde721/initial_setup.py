"""
Initial Setup: Legal agreement document with RECITALS section (no borders)
Task ID: writer_legal_042
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_042'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Title --
    title = doc.add_heading('SOFTWARE DEVELOPMENT AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- Preamble --
    preamble = doc.add_paragraph()
    preamble.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = preamble.add_run(
        'This Software Development Agreement (the "Agreement") is entered into as of '
        'March 15, 2025, by and between Meridian Technologies Inc., a Delaware corporation '
        'with its principal offices at 2400 Innovation Boulevard, Suite 800, San Jose, '
        'California 95134 ("Client"), and Cascade Software Solutions LLC, a California '
        'limited liability company with its principal offices at 1750 Market Street, '
        'Floor 12, San Francisco, California 94102 ("Developer").'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # -- RECITALS heading (centered, bold) --
    recitals_heading = doc.add_heading('RECITALS', level=1)
    recitals_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- WHEREAS clauses (3 paragraphs, NO borders) --
    whereas_texts = [
        'WHEREAS, Client desires to engage Developer to design, develop, and deliver '
        'a cloud-based enterprise resource planning platform (the "Software") capable of '
        'integrating with Client\'s existing SAP and Oracle infrastructure, processing '
        'a minimum of 50,000 concurrent transactions per second, and complying with '
        'SOC 2 Type II and ISO 27001 security standards;',

        'WHEREAS, Developer represents and warrants that it possesses the requisite '
        'technical expertise, qualified personnel, and industry experience to perform '
        'the development services described herein, having successfully delivered similar '
        'enterprise-grade solutions to Fortune 500 clients including Northwind Partners, '
        'Belmont Financial Group, and Pinnacle Healthcare Systems;',

        'WHEREAS, the parties wish to set forth the terms and conditions under which '
        'Developer shall perform the development services, including project milestones, '
        'acceptance criteria, intellectual property ownership, confidentiality obligations, '
        'indemnification provisions, and the compensation structure for the engagement;',
    ]

    for text in whereas_texts:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # -- NOW THEREFORE clause --
    now_therefore = doc.add_paragraph()
    now_therefore.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = now_therefore.add_run(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements herein '
        'contained, and for other good and valuable consideration, the receipt and sufficiency '
        'of which are hereby acknowledged, the parties agree as follows:'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # -- Article I: Scope of Work --
    doc.add_heading('ARTICLE I: SCOPE OF WORK', level=1)

    art1_texts = [
        '1.1  Developer shall design, develop, test, and deploy the Software in accordance '
        'with the functional and technical specifications set forth in Exhibit A attached hereto '
        'and incorporated herein by reference (the "Specifications"). The Software shall include '
        'modules for financial reporting, human resources management, supply chain optimization, '
        'and customer relationship analytics.',

        '1.2  Developer shall assign a dedicated project team consisting of no fewer than eight '
        '(8) senior software engineers, two (2) quality assurance specialists, one (1) DevOps '
        'architect, and one (1) project manager. The project manager shall serve as the primary '
        'point of contact for all communications between the parties.',

        '1.3  Developer shall deliver the Software in three (3) phases as outlined in the '
        'milestone schedule set forth in Exhibit B. Phase 1 (Core Platform) shall be delivered '
        'within ninety (90) calendar days of the Effective Date. Phase 2 (Integration Layer) '
        'shall be delivered within one hundred eighty (180) calendar days. Phase 3 (Analytics '
        'Dashboard) shall be delivered within two hundred seventy (270) calendar days.',
    ]

    for text in art1_texts:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # -- Article II: Compensation --
    doc.add_heading('ARTICLE II: COMPENSATION', level=1)

    art2_texts = [
        '2.1  In consideration for the development services, Client shall pay Developer a '
        'total fixed fee of Eight Hundred Fifty Thousand Dollars ($850,000.00), payable in '
        'installments as follows: (a) $200,000.00 upon execution of this Agreement; '
        '(b) $250,000.00 upon successful completion and acceptance of Phase 1; '
        '(c) $200,000.00 upon successful completion and acceptance of Phase 2; and '
        '(d) $200,000.00 upon successful completion and acceptance of Phase 3.',

        '2.2  All payments shall be made within thirty (30) calendar days of Developer\'s '
        'submission of a proper invoice, accompanied by deliverable documentation and '
        'acceptance test results. Late payments shall accrue interest at the rate of one '
        'and one-half percent (1.5%) per month or the maximum rate permitted by applicable law, '
        'whichever is less.',
    ]

    for text in art2_texts:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # -- Article III: Intellectual Property --
    doc.add_heading('ARTICLE III: INTELLECTUAL PROPERTY', level=1)

    art3_text = (
        '3.1  All intellectual property rights in the Software, including all source code, '
        'object code, documentation, designs, algorithms, and derivative works, shall be '
        'the sole and exclusive property of Client upon full payment of the compensation '
        'set forth in Article II. Developer hereby assigns to Client all right, title, and '
        'interest in and to the Software, including all patent, copyright, trade secret, '
        'and other intellectual property rights therein.'
    )
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(art3_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
