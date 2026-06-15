"""
Initial Setup: Client Directory with US phone numbers in (XXX) XXX-XXXX format
Task ID: writer_frd_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Westbrook & Associates Client Directory', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Confidential - Internal Use Only')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # spacer

    # Intro paragraph
    intro = doc.add_paragraph()
    run = intro.add_run(
        'This directory contains contact information for all active clients '
        'of Westbrook & Associates as of March 2025. Please ensure all phone '
        'numbers are kept up to date and report any changes to the office manager.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Client data: 15 entries with phone numbers in (XXX) XXX-XXXX format
    clients = [
        {
            'name': 'Sarah Chen',
            'company': 'Pacific Rim Technologies',
            'address': '1420 Market Street, Suite 300, San Francisco, CA 94103',
            'phone': '(415) 782-3491',
            'email': 'schen@pacificrimtech.com',
        },
        {
            'name': 'Marcus Johnson',
            'company': 'Elevate Capital Group',
            'address': '250 Park Avenue, 7th Floor, New York, NY 10177',
            'phone': '(212) 554-8762',
            'email': 'mjohnson@elevatecapital.com',
        },
        {
            'name': 'Priya Ramirez',
            'company': 'Horizon Healthcare Solutions',
            'address': '8900 Wilshire Boulevard, Beverly Hills, CA 90211',
            'phone': '(310) 629-4185',
            'email': 'pramirez@horizonhealth.com',
        },
        {
            'name': 'David Nakamura',
            'company': 'Cascade Engineering Inc.',
            'address': '5200 NE Sandy Blvd, Portland, OR 97213',
            'phone': '(503) 841-7623',
            'email': 'dnakamura@cascadeeng.com',
        },
        {
            'name': 'Emily Thornton',
            'company': 'Beacon Financial Advisors',
            'address': '175 Federal Street, Boston, MA 02110',
            'phone': '(617) 392-5847',
            'email': 'ethornton@beaconfin.com',
        },
        {
            'name': 'Carlos Gutierrez',
            'company': 'SunValley Organics LLC',
            'address': '3100 E Camelback Road, Phoenix, AZ 85016',
            'phone': '(602) 718-2934',
            'email': 'cgutierrez@sunvalleyorg.com',
        },
        {
            'name': 'Amanda Li',
            'company': 'Vertex Data Systems',
            'address': '2001 Ross Avenue, Suite 700, Dallas, TX 75201',
            'phone': '(214) 463-8521',
            'email': 'ali@vertexdata.com',
        },
        {
            'name': 'Robert Fitzgerald',
            'company': 'Lakeside Property Management',
            'address': '680 N Lake Shore Drive, Chicago, IL 60611',
            'phone': '(312) 957-3146',
            'email': 'rfitzgerald@lakesidepm.com',
        },
        {
            'name': 'Keiko Yamamoto',
            'company': 'BluePeak Consulting',
            'address': '1900 Ninth Avenue, Seattle, WA 98101',
            'phone': '(206) 284-6739',
            'email': 'kyamamoto@bluepeak.com',
        },
        {
            'name': 'Thomas Okonkwo',
            'company': 'Atlas Logistics Corp.',
            'address': '400 Poydras Street, New Orleans, LA 70130',
            'phone': '(504) 631-9258',
            'email': 'tokonkwo@atlaslogistics.com',
        },
        {
            'name': 'Rachel Morrison',
            'company': 'Sterling Architecture Studio',
            'address': '1010 Common Street, Suite 2400, New Orleans, LA 70112',
            'phone': '(504) 847-3612',
            'email': 'rmorrison@sterlingarch.com',
        },
        {
            'name': 'Hector Vasquez',
            'company': 'Meridian Software Solutions',
            'address': '7700 Windrose Avenue, Plano, TX 75024',
            'phone': '(469) 215-8473',
            'email': 'hvasquez@meridiansw.com',
        },
        {
            'name': 'Diana Kowalski',
            'company': 'Northern Trust Realty',
            'address': '2500 Woodward Avenue, Detroit, MI 48201',
            'phone': '(313) 482-7196',
            'email': 'dkowalski@northerntrust.com',
        },
        {
            'name': 'Anil Patel',
            'company': 'Greenfield BioSciences',
            'address': '9500 Gilman Drive, La Jolla, CA 92093',
            'phone': '(858) 374-5821',
            'email': 'apatel@greenfieldbs.com',
        },
        {
            'name': 'Jessica Moreau',
            'company': 'Summit Legal Partners',
            'address': '1801 K Street NW, Washington, DC 20006',
            'phone': '(202) 596-4138',
            'email': 'jmoreau@summitlegal.com',
        },
    ]

    for i, client in enumerate(clients):
        # Section heading with client name
        heading = doc.add_heading(f'{i + 1}. {client["name"]}', level=2)

        # Company
        p = doc.add_paragraph()
        run_label = p.add_run('Company: ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_val = p.add_run(client['company'])
        run_val.font.size = Pt(11)

        # Address
        p = doc.add_paragraph()
        run_label = p.add_run('Address: ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_val = p.add_run(client['address'])
        run_val.font.size = Pt(11)

        # Phone
        p = doc.add_paragraph()
        run_label = p.add_run('Phone: ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_val = p.add_run(client['phone'])
        run_val.font.size = Pt(11)

        # Email
        p = doc.add_paragraph()
        run_label = p.add_run('Email: ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_val = p.add_run(client['email'])
        run_val.font.size = Pt(11)

        # Separator line (except after last client)
        if i < len(clients) - 1:
            sep = doc.add_paragraph()
            sep_run = sep.add_run('_' * 60)
            sep_run.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
            sep_run.font.size = Pt(8)

    # Footer note
    doc.add_paragraph()
    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_note.add_run('Last Updated: March 15, 2025')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
