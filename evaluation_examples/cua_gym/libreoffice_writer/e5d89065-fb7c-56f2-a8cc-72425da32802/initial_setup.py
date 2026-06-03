"""
Initial Setup: Configure chapter numbering for Heading 1 (Roman) and Heading 2 (Arabic subordinate)
Task ID: writer_lec_010
Domain: libreoffice_writer

Creates a Writer document with 3 chapters (Heading 1) each containing 2-4 sections (Heading 2).
No numbering is applied - headings are plain text only.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Document title --
    title = doc.add_heading("Annual Strategic Planning Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        "This document outlines the strategic initiatives, market analysis, "
        "and operational objectives for the upcoming fiscal year. Each chapter "
        "covers a key area of the organization's growth strategy."
    )
    intro.paragraph_format.space_after = Pt(12)

    # =========================================================
    # Chapter 1: Market Analysis (Heading 1, 3 sections Heading 2)
    # =========================================================
    doc.add_heading("Market Analysis", level=1)

    doc.add_paragraph(
        "The following sections provide a comprehensive overview of current market "
        "conditions, competitive landscape, and emerging opportunities identified "
        "during the Q4 2025 review period."
    )

    # Section 1.1
    doc.add_heading("Industry Overview", level=2)
    doc.add_paragraph(
        "The global technology services market reached $1.2 trillion in 2025, "
        "representing a 7.3% year-over-year growth rate. Cloud infrastructure "
        "spending alone accounted for $380 billion, driven by enterprise digital "
        "transformation initiatives across financial services, healthcare, and "
        "manufacturing sectors."
    )
    doc.add_paragraph(
        "Emerging markets in Southeast Asia and Latin America showed particularly "
        "strong growth trajectories, with compound annual growth rates exceeding "
        "12% in managed services adoption."
    )

    # Section 1.2
    doc.add_heading("Competitive Landscape", level=2)
    doc.add_paragraph(
        "Our primary competitors have shifted their strategies toward integrated "
        "platform offerings. Nextera Solutions expanded into AI-driven analytics, "
        "while Meridian Corp acquired two cybersecurity startups to strengthen "
        "their managed detection portfolio. We must respond with differentiated "
        "service bundles targeting mid-market enterprises."
    )

    # Section 1.3
    doc.add_heading("Customer Segmentation", level=2)
    doc.add_paragraph(
        "Analysis of our current customer base reveals three primary segments: "
        "enterprise accounts generating $500K+ annual revenue (18% of clients, "
        "62% of revenue), mid-market accounts at $100K-$500K (34% of clients, "
        "28% of revenue), and SMB accounts under $100K (48% of clients, 10% of "
        "revenue). Strategic focus should prioritize mid-market growth."
    )

    # =========================================================
    # Chapter 2: Financial Projections (Heading 1, 4 sections Heading 2)
    # =========================================================
    doc.add_heading("Financial Projections", level=1)

    doc.add_paragraph(
        "Based on current performance metrics and market conditions, the finance "
        "team has prepared the following projections for FY2026. These figures "
        "assume moderate economic growth and stable exchange rates."
    )

    # Section 2.1
    doc.add_heading("Revenue Forecast", level=2)
    doc.add_paragraph(
        "Total projected revenue for FY2026 is $45.2 million, representing a "
        "15% increase over FY2025 actual revenue of $39.3 million. Subscription "
        "revenue is expected to grow from $24.1 million to $29.8 million as "
        "existing customers migrate to annual licensing models."
    )

    # Section 2.2
    doc.add_heading("Operating Expenses", level=2)
    doc.add_paragraph(
        "Operating expenses are projected at $38.6 million, including $18.2 million "
        "in personnel costs (reflecting 25 new hires across engineering and sales), "
        "$8.4 million in infrastructure and hosting, and $6.1 million in sales and "
        "marketing. Research and development investment increases to $5.9 million."
    )

    # Section 2.3
    doc.add_heading("Capital Expenditure", level=2)
    doc.add_paragraph(
        "Planned capital expenditures total $4.8 million, primarily allocated to "
        "data center expansion in Frankfurt ($2.1 million), office renovation at "
        "the Portland headquarters ($1.4 million), and new development tooling "
        "and equipment ($1.3 million)."
    )

    # Section 2.4
    doc.add_heading("Profitability Targets", level=2)
    doc.add_paragraph(
        "EBITDA margin is targeted at 18.5%, up from 15.2% in FY2025. Net income "
        "is projected at $4.8 million after tax, with free cash flow of $6.2 million. "
        "The board has approved a dividend increase to $0.45 per share, reflecting "
        "confidence in sustainable growth."
    )

    # =========================================================
    # Chapter 3: Operational Strategy (Heading 1, 2 sections Heading 2)
    # =========================================================
    doc.add_heading("Operational Strategy", level=1)

    doc.add_paragraph(
        "Operational excellence remains a cornerstone of our competitive advantage. "
        "The initiatives outlined below focus on improving delivery efficiency, "
        "customer satisfaction, and talent retention."
    )

    # Section 3.1
    doc.add_heading("Process Optimization", level=2)
    doc.add_paragraph(
        "The operations team will implement a revised project management framework "
        "based on hybrid Agile-Waterfall methodology. Key metrics include reducing "
        "average project delivery time from 14 weeks to 11 weeks, improving on-time "
        "delivery rate from 78% to 90%, and decreasing post-launch defect rates by "
        "40% through enhanced QA automation."
    )

    # Section 3.2
    doc.add_heading("Talent Acquisition and Retention", level=2)
    doc.add_paragraph(
        "Human Resources will focus on reducing voluntary turnover from 22% to 15% "
        "through competitive compensation adjustments, expanded remote work policies, "
        "and a new professional development stipend of $3,000 per employee annually. "
        "Targeted hiring will add 15 senior engineers, 6 solution architects, and "
        "4 enterprise sales representatives during Q1-Q2 2026."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
