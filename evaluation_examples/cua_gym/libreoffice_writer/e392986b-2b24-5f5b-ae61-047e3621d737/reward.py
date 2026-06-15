"""
Reward Script: Place top and bottom border on abstract paragraph of academic paper.
Task ID: writer_para_031
Domain: libreoffice_writer
Scoring:
  Component 1: Abstract paragraph has a top border (single line style)  — 0.4 pts
  Component 2: Abstract paragraph has a bottom border (single line style) — 0.4 pts
  Component 3: Both borders use gray color #808080 and sz=2 (0.25pt width) — 0.2 pts
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_para_031'
FILE_PATH = f'{WORKDIR}/Desktop/academic_paper.docx'

# Namespace URI for Word XML
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def find_abstract_paragraph(doc):
    """
    Find the abstract body paragraph: the paragraph immediately following
    the 'Abstract' Heading 2 paragraph.
    Returns the paragraph object, or None if not found.
    """
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 2' and 'abstract' in para.text.lower():
            # The abstract body is the next paragraph
            if i + 1 < len(doc.paragraphs):
                return doc.paragraphs[i + 1]
    return None


def get_paragraph_borders(para):
    """
    Extract paragraph borders from the XML pBdr element.
    Returns a dict of {border_type: {attr_name: attr_value}}.
    """
    borders = {}
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return borders
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        return borders
    for child in pBdr:
        # tag is like {namespace}top, {namespace}bottom, etc.
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        attrs = {}
        for attr_key, attr_val in child.attrib.items():
            # Normalize attribute keys by stripping namespace
            short_key = attr_key.split('}')[1] if '}' in attr_key else attr_key
            attrs[short_key] = attr_val
        borders[tag] = attrs
    return borders


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate abstract paragraph (precondition gate)
    abstract_para = find_abstract_paragraph(doc)
    if abstract_para is None:
        print("CRITICAL: Could not find abstract paragraph (paragraph after 'Abstract' Heading 2)")
        print("REWARD: 0.0")
        return 0.0

    abstract_text_preview = abstract_para.text[:60]
    print(f"INFO: Found abstract paragraph: '{abstract_text_preview}...'")

    # Extract borders
    borders = get_paragraph_borders(abstract_para)
    print(f"INFO: Paragraph borders found: {list(borders.keys())}")

    # Component 1: Top border exists with 'single' line style (0.4 points)
    try:
        top_border = borders.get('top', {})
        top_val = top_border.get('val', '')
        if top_val == 'single':
            print(f"PASS: Component 1 — Top border has style 'single' (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected top border with val='single', found: val={top_val!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Bottom border exists with 'single' line style (0.4 points)
    try:
        bottom_border = borders.get('bottom', {})
        bottom_val = bottom_border.get('val', '')
        if bottom_val == 'single':
            print(f"PASS: Component 2 — Bottom border has style 'single' (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Expected bottom border with val='single', found: val={bottom_val!r}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Both borders use color #808080 and sz=2 (0.25pt = 2 eighths of a point) (0.2 points)
    try:
        top_border = borders.get('top', {})
        bottom_border = borders.get('bottom', {})

        top_color = top_border.get('color', '').upper()
        bottom_color = bottom_border.get('color', '').upper()
        top_sz = top_border.get('sz', '')
        bottom_sz = bottom_border.get('sz', '')

        color_ok = (top_color == '808080' and bottom_color == '808080')
        # sz=2 corresponds to 2 eighths of a point = 0.25pt
        sz_ok = (top_sz == '2' and bottom_sz == '2')

        if color_ok and sz_ok:
            print(f"PASS: Component 3 — Both borders have color=#808080 and sz=2 (0.25pt) (0.2 pts)")
            total_score += 0.2
        else:
            if not color_ok:
                print(f"FAIL: Component 3 — Color check failed: top_color={top_color!r}, bottom_color={bottom_color!r}, expected '808080'")
            if not sz_ok:
                print(f"FAIL: Component 3 — Size check failed: top_sz={top_sz!r}, bottom_sz={bottom_sz!r}, expected '2' (0.25pt)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
