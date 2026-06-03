"""
Initial Setup: Academic paper docx with abstract paragraph (no borders yet)
Task ID: writer_para_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'academic_paper'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Paragraph 1: Title - Heading 1
    doc.add_heading(
        'Exploring the Relationship Between Sleep Quality and Academic Performance Among University Students',
        level=1
    )

    # Paragraph 2: Authors - normal paragraph
    authors_para = doc.add_paragraph('Authors: Dr. Lisa Park, Prof. James Wright, Dr. Arun Patel')

    # Paragraph 3: 'Abstract' - Heading 2
    doc.add_heading('Abstract', level=2)

    # Paragraph 4: Abstract body - no border (task is to add it)
    abstract_text = (
        'This cross-sectional study examined the association between self-reported sleep quality '
        'and grade point average among 2,340 undergraduate students at a large public university. '
        'Using the Pittsburgh Sleep Quality Index and institutional academic records, we found a '
        'statistically significant positive correlation between sleep quality scores and cumulative GPA. '
        'Students reporting poor sleep quality had GPAs averaging 0.4 points lower than their '
        'well-rested peers.'
    )
    abstract_para = doc.add_paragraph(abstract_text)

    # Paragraph 5: Keywords - normal paragraph
    doc.add_paragraph(
        'Keywords: sleep quality, academic performance, university students, PSQI, GPA'
    )

    # Paragraph 6: Introduction heading
    doc.add_heading('1. Introduction', level=2)

    # Additional introduction content to make document realistic
    doc.add_paragraph(
        'Sleep is a fundamental biological process that plays a critical role in cognitive functioning, '
        'memory consolidation, and overall health. For university students, maintaining adequate sleep '
        'quality can be particularly challenging due to academic demands, social activities, and '
        'irregular schedules. Despite growing awareness of sleep-related issues among college populations, '
        'empirical data on the direct relationship between sleep quality and academic outcomes remains '
        'an active area of investigation.'
    )

    doc.add_heading('2. Methods', level=2)

    doc.add_paragraph(
        'Participants were recruited from a large public university in the southeastern United States. '
        'Inclusion criteria required enrollment as a full-time undergraduate student with a minimum of '
        'one completed academic semester. The Pittsburgh Sleep Quality Index (PSQI) was administered '
        'electronically during the fall semester of 2024. Cumulative GPA data were obtained from '
        'institutional academic records with appropriate consent.'
    )

    doc.add_heading('3. Results', level=2)

    doc.add_paragraph(
        'A total of 2,340 students completed the PSQI survey (response rate: 78.4%). The mean PSQI '
        'global score was 6.8 (SD = 3.2), with 54.2% of students classified as poor sleepers '
        '(PSQI > 5). Pearson correlation analysis revealed a significant negative correlation between '
        'PSQI score and GPA (r = -0.42, p < 0.001), indicating that higher PSQI scores (worse sleep) '
        'were associated with lower academic performance.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open with LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
