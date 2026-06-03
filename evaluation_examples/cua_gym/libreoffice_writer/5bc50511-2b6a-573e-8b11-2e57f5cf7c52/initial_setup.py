"""
Initial Setup: Legal contract document with continuous Arabic page numbering
Task ID: writer_legal_065
Domain: libreoffice_writer

Creates a Writer document with:
- Cover page
- Table of Contents placeholder
- 20 pages of contract body
- 8 pages of exhibits
All using continuous Arabic numbering (single default section).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_065'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add an explicit page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Add simple page number footer (continuous Arabic for all pages)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # PAGE field code
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)

    # ===== COVER PAGE (page 1) =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading("", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("MASTER SERVICES AGREEMENT")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x36, 0x64)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Between")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()

    party1 = doc.add_paragraph()
    party1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = party1.add_run("Meridian Technology Solutions, Inc.")
    run.font.size = Pt(18)
    run.bold = True

    and_p = doc.add_paragraph()
    and_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = and_p.add_run("and")
    run.font.size = Pt(14)
    run.font.italic = True

    party2 = doc.add_paragraph()
    party2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = party2.add_run("Pinnacle Global Enterprises, LLC")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_p.add_run("Effective Date: March 15, 2025")
    run.font.size = Pt(12)

    ref_p = doc.add_paragraph()
    ref_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = ref_p.add_run("Agreement No. MSA-2025-04782")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ===== TABLE OF CONTENTS (page 2) =====
    add_page_break(doc)

    toc_title = doc.add_heading("TABLE OF CONTENTS", level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    toc_entries = [
        ("Article I", "Definitions and Interpretation", "1"),
        ("Article II", "Scope of Services", "3"),
        ("Article III", "Term and Termination", "5"),
        ("Article IV", "Compensation and Payment Terms", "7"),
        ("Article V", "Intellectual Property Rights", "9"),
        ("Article VI", "Confidentiality", "10"),
        ("Article VII", "Representations and Warranties", "12"),
        ("Article VIII", "Indemnification", "14"),
        ("Article IX", "Limitation of Liability", "15"),
        ("Article X", "Dispute Resolution", "16"),
        ("Article XI", "Force Majeure", "17"),
        ("Article XII", "General Provisions", "18"),
        ("Exhibit A", "Statement of Work", "21"),
        ("Exhibit B", "Service Level Agreement", "23"),
        ("Exhibit C", "Data Processing Addendum", "25"),
        ("Exhibit D", "Insurance Requirements", "27"),
    ]
    for art, title_text, pg in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(f"{art}  {title_text}")
        run.font.size = Pt(11)
        run2 = p.add_run(f"\t{pg}")
        run2.font.size = Pt(11)

    # ===== CONTRACT BODY — Articles I through XII (~20 pages) =====
    articles = [
        ("Article I: Definitions and Interpretation", [
            '1.1 "Affiliate" means any entity that directly or indirectly controls, is controlled by, or is under common control with a party to this Agreement, where "control" means the ownership of more than fifty percent (50%) of the voting securities or equivalent ownership interest.',
            '1.2 "Confidential Information" means all non-public information disclosed by either party to the other, whether orally, in writing, or by inspection, including but not limited to business plans, financial data, customer lists, technical specifications, trade secrets, and proprietary algorithms.',
            '1.3 "Deliverables" means all work product, documentation, software, reports, and other materials created, developed, or produced by Service Provider in the course of performing the Services under this Agreement.',
            '1.4 "Effective Date" means March 15, 2025, or such later date as the parties may agree in writing.',
            '1.5 "Force Majeure Event" means any event beyond the reasonable control of the affected party, including but not limited to acts of God, natural disasters, epidemics, pandemics, government actions, war, terrorism, labor disputes, and failures of third-party telecommunications or power supply.',
            '1.6 "Intellectual Property" means all patents, copyrights, trademarks, trade secrets, know-how, and any other intellectual property rights recognized in any jurisdiction worldwide.',
            '1.7 "Service Level Agreement" or "SLA" means the performance standards and metrics set forth in Exhibit B, as may be amended from time to time by mutual written agreement of the parties.',
            '1.8 "Services" means the consulting, development, implementation, and support services described in the applicable Statement of Work attached as Exhibit A.',
        ]),
        ("Article II: Scope of Services", [
            "2.1 Service Provider agrees to perform the Services described in each Statement of Work executed by the parties. Each Statement of Work shall be substantially in the form attached hereto as Exhibit A and shall be incorporated into and governed by the terms of this Agreement.",
            "2.2 Service Provider shall assign qualified personnel with appropriate skills and experience to perform the Services. Client shall have the right to request replacement of any Service Provider personnel who, in Client's reasonable judgment, are not performing satisfactorily.",
            "2.3 Service Provider shall perform the Services in a professional and workmanlike manner, consistent with generally accepted industry standards and practices. Service Provider shall comply with all applicable laws, regulations, and Client's reasonable policies communicated in writing.",
            "2.4 Any changes to the scope of Services shall be documented in a written change order signed by authorized representatives of both parties. No oral modifications shall be binding.",
            "2.5 Service Provider shall maintain detailed records of all hours worked, expenses incurred, and Deliverables produced in connection with the Services, and shall make such records available to Client upon reasonable request.",
        ]),
        ("Article III: Term and Termination", [
            "3.1 This Agreement shall commence on the Effective Date and continue for an initial term of three (3) years, unless earlier terminated in accordance with this Article III (the 'Initial Term').",
            "3.2 Upon expiration of the Initial Term, this Agreement shall automatically renew for successive one (1) year periods (each a 'Renewal Term'), unless either party provides written notice of non-renewal at least ninety (90) days prior to the expiration of the then-current term.",
            "3.3 Either party may terminate this Agreement for cause upon sixty (60) days' prior written notice to the other party if the other party commits a material breach of this Agreement and fails to cure such breach within sixty (60) days after receipt of written notice specifying the breach.",
            "3.4 Client may terminate this Agreement for convenience upon one hundred twenty (120) days' prior written notice to Service Provider, subject to payment of all fees earned through the effective date of termination and any applicable early termination fees specified in the relevant Statement of Work.",
            "3.5 Upon termination or expiration of this Agreement, Service Provider shall promptly return or destroy all Confidential Information of Client and provide reasonable transition assistance for a period not to exceed ninety (90) days.",
        ]),
        ("Article IV: Compensation and Payment Terms", [
            "4.1 Client shall pay Service Provider the fees set forth in the applicable Statement of Work. Unless otherwise specified, fees shall be invoiced monthly in arrears based on actual hours worked at the rates specified in the Statement of Work.",
            "4.2 All invoices shall be payable within thirty (30) days of receipt. Late payments shall bear interest at the lesser of one and one-half percent (1.5%) per month or the maximum rate permitted by applicable law.",
            "4.3 Service Provider shall be reimbursed for reasonable and pre-approved out-of-pocket expenses incurred in connection with the Services, provided that any individual expense exceeding Five Hundred Dollars ($500.00) requires prior written approval from Client.",
            "4.4 All fees are exclusive of applicable taxes. Client shall be responsible for all sales, use, value-added, withholding, and similar taxes arising from the transactions contemplated by this Agreement, excluding taxes based on Service Provider's income.",
            "4.5 Client shall have the right to audit Service Provider's time records and expense reports upon reasonable notice, no more than once per calendar year, during regular business hours.",
        ]),
        ("Article V: Intellectual Property Rights", [
            "5.1 All Deliverables created by Service Provider specifically for Client under this Agreement shall be considered 'works made for hire' to the maximum extent permitted by applicable law. To the extent any Deliverable does not qualify as a work made for hire, Service Provider hereby irrevocably assigns to Client all right, title, and interest in and to such Deliverable.",
            "5.2 Service Provider retains all rights in its pre-existing intellectual property, tools, methodologies, and know-how ('Service Provider IP'). To the extent any Service Provider IP is incorporated into a Deliverable, Service Provider grants Client a perpetual, non-exclusive, royalty-free license to use such Service Provider IP solely as part of the Deliverable.",
            "5.3 Client retains all rights in its pre-existing intellectual property, data, and materials provided to Service Provider for use in performing the Services ('Client IP'). Service Provider shall use Client IP solely for the purpose of performing the Services.",
            "5.4 Neither party shall use the other party's trademarks, trade names, or logos without the prior written consent of the other party, except as reasonably necessary to perform obligations under this Agreement.",
        ]),
        ("Article VI: Confidentiality", [
            "6.1 Each party agrees to hold in strict confidence all Confidential Information received from the other party and shall not disclose such Confidential Information to any third party without the prior written consent of the disclosing party, except to employees, contractors, and advisors who have a need to know and are bound by confidentiality obligations at least as protective as those contained herein.",
            "6.2 The obligations of confidentiality shall not apply to information that: (a) is or becomes publicly available through no fault of the receiving party; (b) was already known to the receiving party prior to disclosure; (c) is independently developed by the receiving party without reference to the Confidential Information; or (d) is disclosed pursuant to a valid court order or governmental regulation, provided the receiving party gives prompt notice to the disclosing party.",
            "6.3 The confidentiality obligations under this Article VI shall survive the termination or expiration of this Agreement for a period of five (5) years, or with respect to trade secrets, for so long as such information remains a trade secret under applicable law.",
            "6.4 Each party acknowledges that any breach of the confidentiality obligations may cause irreparable harm and that the non-breaching party shall be entitled to seek equitable relief, including injunction and specific performance, in addition to all other remedies available at law or in equity.",
        ]),
        ("Article VII: Representations and Warranties", [
            "7.1 Service Provider represents and warrants that: (a) it has the legal right and authority to enter into this Agreement; (b) the Services will be performed in a professional manner consistent with industry standards; (c) the Deliverables will conform to the specifications set forth in the applicable Statement of Work; and (d) the Deliverables will not infringe any third party's intellectual property rights.",
            "7.2 Client represents and warrants that: (a) it has the legal right and authority to enter into this Agreement; (b) it has all necessary rights to provide Client IP to Service Provider for use in performing the Services; and (c) it will cooperate reasonably with Service Provider to facilitate performance of the Services.",
            "7.3 EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, NEITHER PARTY MAKES ANY WARRANTIES, EXPRESS OR IMPLIED, INCLUDING WITHOUT LIMITATION ANY IMPLIED WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, OR NON-INFRINGEMENT.",
        ]),
        ("Article VIII: Indemnification", [
            "8.1 Service Provider shall indemnify, defend, and hold harmless Client and its officers, directors, employees, and agents from and against any and all third-party claims, losses, damages, liabilities, costs, and expenses (including reasonable attorneys' fees) arising from or related to: (a) Service Provider's negligence or willful misconduct; (b) any breach of Service Provider's representations, warranties, or obligations under this Agreement; or (c) any claim that the Deliverables infringe a third party's intellectual property rights.",
            "8.2 Client shall indemnify, defend, and hold harmless Service Provider and its officers, directors, employees, and agents from and against any and all third-party claims, losses, damages, liabilities, costs, and expenses (including reasonable attorneys' fees) arising from or related to: (a) Client's negligence or willful misconduct; (b) any breach of Client's representations, warranties, or obligations under this Agreement; or (c) any claim that Client IP infringes a third party's intellectual property rights.",
            "8.3 The indemnifying party's obligations under this Article VIII are conditioned upon: (a) prompt written notice of the claim; (b) sole control of the defense and settlement; and (c) reasonable cooperation from the indemnified party at the indemnifying party's expense.",
        ]),
        ("Article IX: Limitation of Liability", [
            "9.1 IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER PARTY FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING BUT NOT LIMITED TO LOSS OF PROFITS, LOSS OF DATA, BUSINESS INTERRUPTION, OR LOSS OF GOODWILL, REGARDLESS OF THE CAUSE OF ACTION OR THE THEORY OF LIABILITY, EVEN IF SUCH PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.",
            "9.2 THE TOTAL AGGREGATE LIABILITY OF EITHER PARTY UNDER THIS AGREEMENT SHALL NOT EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT TO SERVICE PROVIDER DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO THE CLAIM.",
            "9.3 The limitations of liability set forth in this Article IX shall not apply to: (a) breaches of confidentiality obligations under Article VI; (b) indemnification obligations under Article VIII; (c) infringement of intellectual property rights; or (d) either party's gross negligence or willful misconduct.",
        ]),
        ("Article X: Dispute Resolution", [
            "10.1 The parties shall attempt in good faith to resolve any dispute arising out of or relating to this Agreement through informal negotiation between senior executives of each party for a period of thirty (30) days following written notice of the dispute.",
            "10.2 If the dispute is not resolved through informal negotiation, the parties agree to submit the dispute to binding arbitration administered by the American Arbitration Association under its Commercial Arbitration Rules. The arbitration shall be conducted by a panel of three (3) arbitrators in New York, New York.",
            "10.3 The arbitrators shall have the authority to award any remedy available at law or in equity, including injunctive relief and specific performance. The arbitral award shall be final and binding and may be entered as a judgment in any court of competent jurisdiction.",
            "10.4 Notwithstanding the foregoing, either party may seek provisional or injunctive relief from a court of competent jurisdiction to prevent irreparable harm pending the outcome of the arbitration.",
        ]),
        ("Article XI: Force Majeure", [
            "11.1 Neither party shall be liable for any delay or failure to perform its obligations under this Agreement to the extent that such delay or failure is caused by a Force Majeure Event, provided that the affected party gives prompt notice to the other party and uses commercially reasonable efforts to mitigate the effects of the Force Majeure Event.",
            "11.2 If a Force Majeure Event continues for more than ninety (90) consecutive days, either party may terminate this Agreement upon thirty (30) days' written notice to the other party, without liability for such termination other than payment for Services rendered prior to the effective date of termination.",
        ]),
        ("Article XII: General Provisions", [
            "12.1 Governing Law. This Agreement shall be governed by and construed in accordance with the laws of the State of New York, without regard to its conflict of laws principles.",
            "12.2 Entire Agreement. This Agreement, including all Statements of Work and Exhibits attached hereto, constitutes the entire agreement between the parties and supersedes all prior or contemporaneous agreements, negotiations, and discussions, whether oral or written.",
            "12.3 Amendment. This Agreement may not be amended or modified except by a written instrument signed by authorized representatives of both parties.",
            "12.4 Assignment. Neither party may assign or transfer this Agreement or any rights or obligations hereunder without the prior written consent of the other party, except that either party may assign this Agreement to a successor in connection with a merger, acquisition, or sale of substantially all of its assets.",
            "12.5 Notices. All notices under this Agreement shall be in writing and delivered by certified mail, overnight courier, or electronic mail to the addresses set forth on the signature page, or such other address as a party may designate in writing.",
            "12.6 Severability. If any provision of this Agreement is held to be invalid or unenforceable, the remaining provisions shall continue in full force and effect.",
            "12.7 Waiver. The failure of either party to enforce any provision of this Agreement shall not constitute a waiver of that party's right to enforce that provision or any other provision in the future.",
            "12.8 Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original and all of which together shall constitute one and the same instrument.",
        ]),
    ]

    # Write contract body articles
    for article_idx, (article_title, paragraphs) in enumerate(articles):
        if article_idx == 0:
            add_page_break(doc)

        heading = doc.add_heading(article_title, level=1)

        for para_text in paragraphs:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'

        # Add spacing/breaks between articles to fill ~20 pages
        if article_idx < len(articles) - 1:
            doc.add_paragraph()
            # Add page break every 2-3 articles to distribute content
            if article_idx % 2 == 1:
                add_page_break(doc)

    # ===== EXHIBITS (~8 pages) =====
    add_page_break(doc)

    # Exhibit A: Statement of Work
    doc.add_heading("EXHIBIT A: STATEMENT OF WORK", level=1)
    doc.add_paragraph()

    sow_sections = [
        ("1. Project Overview", "Meridian Technology Solutions shall provide comprehensive enterprise software development and integration services for Pinnacle Global Enterprises' next-generation customer relationship management (CRM) platform, including custom module development, third-party API integrations, data migration, and end-user training."),
        ("2. Project Timeline", "Phase 1 - Discovery and Requirements (Weeks 1-4): Conduct stakeholder interviews, document business requirements, and develop detailed technical specifications.\nPhase 2 - Design and Architecture (Weeks 5-8): Create system architecture, database design, and UI/UX mockups.\nPhase 3 - Development (Weeks 9-24): Iterative development using agile methodology with bi-weekly sprint reviews.\nPhase 4 - Testing and QA (Weeks 25-28): Comprehensive testing including unit, integration, performance, and user acceptance testing.\nPhase 5 - Deployment and Go-Live (Weeks 29-30): Staged deployment with rollback procedures and 24/7 support."),
        ("3. Staffing and Rates", "Project Manager: $185/hour\nSenior Developer (x3): $165/hour each\nQA Engineer (x2): $135/hour each\nUX Designer: $155/hour\nDatabase Architect: $175/hour\nDevOps Engineer: $160/hour\nEstimated Total: $1,847,500 (not to exceed without written approval)"),
        ("4. Deliverables", "- Requirements Specification Document\n- System Architecture Document\n- Database Schema and Migration Scripts\n- Custom CRM Modules (Sales, Marketing, Service)\n- API Integration Layer\n- Automated Test Suite\n- Deployment Runbook\n- End-User Training Materials\n- Administrator Guide"),
    ]
    for title, content in sow_sections:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    add_page_break(doc)
    add_page_break(doc)

    # Exhibit B: Service Level Agreement
    doc.add_heading("EXHIBIT B: SERVICE LEVEL AGREEMENT", level=1)
    doc.add_paragraph()

    sla_sections = [
        ("1. Service Availability", "Service Provider guarantees 99.9% uptime for all production systems during business hours (Monday-Friday, 8:00 AM - 8:00 PM EST), excluding scheduled maintenance windows communicated at least 72 hours in advance."),
        ("2. Response Time Standards", "Priority 1 (Critical): Response within 15 minutes, resolution target 4 hours\nPriority 2 (High): Response within 1 hour, resolution target 8 hours\nPriority 3 (Medium): Response within 4 hours, resolution target 2 business days\nPriority 4 (Low): Response within 1 business day, resolution target 5 business days"),
        ("3. Performance Metrics", "Page load time: < 2 seconds (95th percentile)\nAPI response time: < 500ms (99th percentile)\nConcurrent users supported: minimum 5,000\nData processing throughput: minimum 10,000 records per minute"),
        ("4. Service Credits", "If availability falls below the guaranteed threshold:\n99.0% - 99.9%: 5% service credit\n95.0% - 99.0%: 10% service credit\nBelow 95.0%: 25% service credit\nService credits shall be applied to the next monthly invoice."),
    ]
    for title, content in sla_sections:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    add_page_break(doc)
    add_page_break(doc)

    # Exhibit C: Data Processing Addendum
    doc.add_heading("EXHIBIT C: DATA PROCESSING ADDENDUM", level=1)
    doc.add_paragraph()

    dpa_sections = [
        ("1. Data Processing Scope", "Service Provider shall process personal data only as necessary to perform the Services and in accordance with Client's documented instructions. Categories of data subjects include: Client employees, Client customers, and prospective customers. Types of personal data include: names, email addresses, phone numbers, employment records, and transaction histories."),
        ("2. Security Measures", "Service Provider shall implement and maintain appropriate technical and organizational security measures, including:\n- Encryption of data at rest (AES-256) and in transit (TLS 1.3)\n- Multi-factor authentication for all system access\n- Regular security audits and penetration testing\n- Employee background checks and security training\n- Incident response plan with 24-hour notification requirement"),
        ("3. Data Subject Rights", "Service Provider shall assist Client in responding to data subject requests regarding access, rectification, erasure, portability, and objection to processing, within the timeframes required by applicable data protection laws."),
        ("4. Cross-Border Transfers", "Any transfer of personal data outside the European Economic Area shall be subject to appropriate safeguards, including Standard Contractual Clauses (Module 2: Controller to Processor) as approved by the European Commission."),
    ]
    for title, content in dpa_sections:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    add_page_break(doc)
    add_page_break(doc)

    # Exhibit D: Insurance Requirements
    doc.add_heading("EXHIBIT D: INSURANCE REQUIREMENTS", level=1)
    doc.add_paragraph()

    ins_sections = [
        ("1. Required Coverage", "Service Provider shall maintain the following insurance coverages throughout the term of this Agreement:\n- Commercial General Liability: $5,000,000 per occurrence / $10,000,000 aggregate\n- Professional Liability (Errors & Omissions): $5,000,000 per claim / $10,000,000 aggregate\n- Cyber Liability: $10,000,000 per occurrence\n- Workers' Compensation: As required by applicable law\n- Commercial Auto Liability: $1,000,000 combined single limit"),
        ("2. Insurance Standards", "All insurance policies shall be issued by carriers rated A- VII or better by A.M. Best Company. Service Provider shall name Client as an additional insured on the Commercial General Liability and Cyber Liability policies."),
        ("3. Certificates and Notice", "Service Provider shall provide certificates of insurance to Client within ten (10) business days of the Effective Date and upon each policy renewal. Service Provider shall provide Client with at least thirty (30) days' advance written notice of any material change, cancellation, or non-renewal of required coverages."),
    ]
    for title, content in ins_sections:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
