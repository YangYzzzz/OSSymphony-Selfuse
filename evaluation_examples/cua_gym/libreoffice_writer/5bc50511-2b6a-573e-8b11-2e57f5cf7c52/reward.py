"""
Reward Script: Page style sections with Roman, Arabic, and Exhibit numbering
Task ID: writer_legal_065
Domain: libreoffice_writer
Scoring:
  Component 1: Document has 3 sections (0.20)
  Component 2: Front matter section uses lowerRoman numbering starting at 1 (0.20)
  Component 3: Main body section uses decimal numbering starting at 1 (0.20)
  Component 4: Exhibit section uses decimal numbering starting at 1 (0.15)
  Component 5: Exhibit section footer has "Ex-" prefix before page number (0.15)
  Component 6: Section breaks at correct transition points (0.10)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_065'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)

    # Component 1: Document has exactly 3 sections (0.20 points)
    # Initial has 1 section; golden has 3.
    try:
        if num_sections >= 3:
            print(f"PASS: Component 1 — Document has {num_sections} sections (>= 3) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected >= 3 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # If fewer than 3 sections, remaining checks will fail gracefully
    if num_sections < 3:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Helper: extract pgNumType attributes from a section
    def get_page_num_type(section):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            return None, None
        fmt = pgNumType.get(qn('w:fmt'))
        start = pgNumType.get(qn('w:start'))
        return fmt, start

    # Component 2: Section 0 (Front Matter) uses lowerRoman numbering, start=1 (0.20 points)
    try:
        fmt0, start0 = get_page_num_type(doc.sections[0])
        if fmt0 == 'lowerRoman' and start0 is not None and int(start0) == 1:
            print(f"PASS: Component 2 — Section 0 fmt=lowerRoman, start=1 (0.20 pts)")
            total_score += 0.20
        elif fmt0 == 'lowerRoman':
            # lowerRoman format is correct but start might be missing/different
            print(f"PARTIAL: Component 2 — Section 0 fmt=lowerRoman but start={start0} (0.10 pts)")
            if start0 is not None:
                total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Section 0 expected fmt=lowerRoman, found fmt={fmt0}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Section 1 (Main Body) uses decimal numbering, start=1 (0.20 points)
    try:
        fmt1, start1 = get_page_num_type(doc.sections[1])
        if fmt1 == 'decimal' and start1 is not None and int(start1) == 1:
            print(f"PASS: Component 3 — Section 1 fmt=decimal, start=1 (0.20 pts)")
            total_score += 0.20
        elif fmt1 == 'decimal':
            print(f"PARTIAL: Component 3 — Section 1 fmt=decimal but start={start1} (0.10 pts)")
            if start1 is not None:
                total_score += 0.10
        else:
            print(f"FAIL: Component 3 — Section 1 expected fmt=decimal, found fmt={fmt1}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Section 2 (Exhibit) uses decimal numbering, start=1 (0.15 points)
    try:
        fmt2, start2 = get_page_num_type(doc.sections[2])
        if fmt2 == 'decimal' and start2 is not None and int(start2) == 1:
            print(f"PASS: Component 4 — Section 2 fmt=decimal, start=1 (0.15 pts)")
            total_score += 0.15
        elif fmt2 == 'decimal':
            print(f"PARTIAL: Component 4 — Section 2 fmt=decimal but start={start2} (0.07 pts)")
            if start2 is not None:
                total_score += 0.07
        else:
            print(f"FAIL: Component 4 — Section 2 expected fmt=decimal, found fmt={fmt2}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Section 2 footer has "Ex-" prefix text before the page number field (0.15 points)
    try:
        footer = doc.sections[2].footer
        footer_text = ""
        page_field_count = 0
        for para in footer.paragraphs:
            footer_text += para.text
            # Count PAGE field codes
            instr_elements = para._element.findall('.//' + qn('w:instrText'))
            for el in instr_elements:
                if el.text and 'PAGE' in el.text:
                    page_field_count += 1

        # The footer should have "Ex-" prefix text AND a page number field
        if 'Ex-' in footer_text and page_field_count > 0:
            print(f"PASS: Component 5 — Section 2 footer has 'Ex-' prefix and PAGE field (0.15 pts)")
            total_score += 0.15
        elif 'Ex-' in footer_text or 'ex-' in footer_text.lower():
            print(f"PARTIAL: Component 5 — Footer has Ex- text but PAGE fields={page_field_count} (0.07 pts)")
            if page_field_count >= 0:
                total_score += 0.07
        else:
            print(f"FAIL: Component 5 — Section 2 footer text='{footer_text}', page_fields={page_field_count}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Section breaks at correct transition points (0.10 points)
    # The front matter (cover+TOC) ends before main body; main body ends before exhibits.
    # We verify that the section breaks are placed in the right region of the document.
    # In the golden, breaks are at paragraphs [35] and [118] (between TOC and body, and body and exhibits).
    try:
        # Find paragraphs with inline sectPr (section breaks)
        break_indices = []
        for i, para in enumerate(doc.paragraphs):
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                sect = pPr.find(qn('w:sectPr'))
                if sect is not None:
                    break_indices.append(i)

        # We expect at least 2 section breaks to create 3 sections
        # First break should be roughly in the first ~40 paragraphs (end of front matter)
        # Second break should be roughly before the exhibits start (~paragraph 115-120)
        if len(break_indices) >= 2:
            first_break = break_indices[0]
            second_break = break_indices[1]
            # Front matter break should be in first third of doc; exhibit break in last third
            total_paras = len(doc.paragraphs)
            first_ok = first_break < total_paras * 0.4
            second_ok = second_break > total_paras * 0.5
            if first_ok and second_ok:
                print(f"PASS: Component 6 — Section breaks at paras {break_indices} (0.10 pts)")
                total_score += 0.10
            elif first_ok or second_ok:
                print(f"PARTIAL: Component 6 — Breaks at {break_indices}, one position correct (0.05 pts)")
                total_score += 0.05
            else:
                print(f"FAIL: Component 6 — Breaks at {break_indices}, both positions questionable")
        else:
            print(f"FAIL: Component 6 — Expected >= 2 section breaks, found {len(break_indices)} at {break_indices}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
