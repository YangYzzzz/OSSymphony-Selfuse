"""
Initial Setup: Format file paths with character style
Task ID: writer_tech_049
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_049'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    doc.add_heading('Infrastructure Migration Runbook', level=1)

    # Section 1
    doc.add_heading('1. Pre-Migration Checklist', level=2)
    doc.add_paragraph(
        'Before starting the migration process, ensure all configuration files have been '
        'backed up. The primary application configuration is stored at /etc/config.yaml and '
        'should be copied to the staging server before any changes are made. Verify that all '
        'services listed in the configuration are currently running and healthy.'
    )
    doc.add_paragraph(
        'Check the deployment log entries from the past 48 hours to identify any recent '
        'changes that might affect the migration. Coordinate with the networking team to '
        'ensure firewall rules have been updated for the new IP ranges.'
    )

    # Section 2
    doc.add_heading('2. SSH Access Configuration', level=2)
    doc.add_paragraph(
        'All team members who need access to the production servers must have their public '
        'keys added to ~/.ssh/authorized_keys on each target host. The security team has '
        'mandated that password-based authentication be disabled as of Q2 2025. Rotate all '
        'existing keys that are older than 90 days.'
    )
    doc.add_paragraph(
        'For the bastion host setup, configure ProxyJump rules in your local SSH config. '
        'Ensure that the SSH daemon is configured to use only Ed25519 and RSA keys with a '
        'minimum length of 4096 bits.'
    )

    # Section 3
    doc.add_heading('3. Logging and Monitoring', level=2)
    doc.add_paragraph(
        'System logs are centrally collected and can be reviewed at /var/log/syslog on each '
        'node. During the migration window, set up real-time log tailing on all critical '
        'services. The monitoring dashboard should be configured to alert on any error rate '
        'exceeding 2% within a 5-minute window.'
    )
    doc.add_paragraph(
        'Ensure that the log rotation policy is in place and that archived logs are being '
        'shipped to the long-term storage bucket every 24 hours. Disk usage on the logging '
        'partition should not exceed 80% capacity.'
    )

    # Section 4
    doc.add_heading('4. Deployment Procedures', level=2)
    doc.add_paragraph(
        'The automated deployment pipeline is managed by the script at /usr/local/bin/deploy.sh '
        'which handles rolling updates across all application nodes. Before triggering a '
        'deployment, verify that the CI pipeline has passed all tests and that the artifact '
        'repository contains the expected build version.'
    )
    doc.add_paragraph(
        'Rollback procedures should be tested in the staging environment at least once per '
        'quarter. The deployment script supports a --rollback flag that reverts to the '
        'previous release within 60 seconds.'
    )

    # Section 5
    doc.add_heading('5. Environment Variables and Shell Configuration', level=2)
    doc.add_paragraph(
        'Each developer\'s shell environment is customized through ~/.bashrc which should '
        'include the necessary PATH exports, aliases, and credential references for the '
        'development toolkit. Do not store plaintext secrets in shell configuration files; '
        'use the vault integration instead.'
    )
    doc.add_paragraph(
        'Standardize the prompt format across the team to include the current Kubernetes '
        'context and namespace. This helps prevent accidental commands against production '
        'clusters.'
    )

    # Section 6
    doc.add_heading('6. Application Settings', level=2)
    doc.add_paragraph(
        'Runtime configuration for the main application is managed through '
        '/opt/app/settings.json which is mounted as a ConfigMap in the Kubernetes deployment. '
        'Changes to this file require a rolling restart of all pods in the application '
        'namespace. Feature flags and A/B test configurations are also defined here.'
    )
    doc.add_paragraph(
        'The settings file supports environment-specific overrides. Production values take '
        'precedence over defaults, and any sensitive values should be injected from the '
        'secrets manager at pod startup time rather than stored in the ConfigMap directly.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
