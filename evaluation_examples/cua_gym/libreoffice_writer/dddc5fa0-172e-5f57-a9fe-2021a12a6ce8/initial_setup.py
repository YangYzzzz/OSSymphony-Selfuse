"""
Initial Setup: Create branded_doc.docx with separate rectangle shape and logo image on page 1
Task ID: writer_obj_071
Domain: libreoffice_writer

Initial state: Page 1 has a dark blue rectangle shape (17cm x 2.5cm) and a logo image
(2cm x 2cm) placed SEPARATELY near the top - NOT grouped, NOT at exact X:0 Y:0.
"""

import os
import io
import re
import shlex
import struct
import subprocess
import time
import zlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_071'
OUTPUT = f'{WORKDIR}/Desktop/branded_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_png_bytes(width_px, height_px, r, g, b):
    """Create a minimal valid PNG image with a solid color."""
    def png_chunk(chunk_type, data):
        c = chunk_type + data
        crc = zlib.crc32(c) & 0xffffffff
        return struct.pack('>I', len(data)) + c + struct.pack('>I', crc)

    ihdr_data = struct.pack('>IIBBBBB', width_px, height_px, 8, 2, 0, 0, 0)
    ihdr = png_chunk(b'IHDR', ihdr_data)

    raw_rows = b''
    for _ in range(height_px):
        row = b'\x00' + bytes([r, g, b] * width_px)
        raw_rows += row
    compressed = zlib.compress(raw_rows)
    idat = png_chunk(b'IDAT', compressed)
    iend = png_chunk(b'IEND', b'')

    return b'\x89PNG\r\n\x1a\n' + ihdr + idat + iend


def cm_to_emu(cm):
    """Convert centimeters to EMU (English Metric Units)."""
    return int(cm * 360000)


def build_rect_anchor_xml(pos_x_emu, pos_y_emu, width_emu, height_emu, shape_id, name, fill_hex):
    """Build XML for a floating rectangle anchor (wp:anchor with wrapNone, relative to page)."""
    return f'''<w:drawing
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <wp:anchor distT="0" distB="0" distL="0" distR="0"
             simplePos="0" relativeHeight="251658240" behindDoc="0"
             locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page">
      <wp:posOffset>{pos_x_emu}</wp:posOffset>
    </wp:positionH>
    <wp:positionV relativeFrom="page">
      <wp:posOffset>{pos_y_emu}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{shape_id}" name="{name}"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr>
            <a:spLocks noChangeArrowheads="1"/>
          </wps:cNvSpPr>
          <wps:spPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
            <a:solidFill>
              <a:srgbClr val="{fill_hex}"/>
            </a:solidFill>
            <a:ln>
              <a:noFill/>
            </a:ln>
          </wps:spPr>
          <wps:bodyPr/>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''


def build_image_anchor_xml(pos_x_emu, pos_y_emu, width_emu, height_emu, shape_id, name, rid):
    """Build XML for a floating image anchor (wp:anchor with wrapNone, relative to page)."""
    return f'''<w:drawing
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
  <wp:anchor distT="0" distB="0" distL="0" distR="0"
             simplePos="0" relativeHeight="251723776" behindDoc="0"
             locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page">
      <wp:posOffset>{pos_x_emu}</wp:posOffset>
    </wp:positionH>
    <wp:positionV relativeFrom="page">
      <wp:posOffset>{pos_y_emu}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{shape_id}" name="{name}" descr="Logo Image"/>
    <wp:cNvGraphicFramePr>
      <a:graphicFrameLocks noChangeAspect="1"/>
    </wp:cNvGraphicFramePr>
    <a:graphic>
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic>
          <pic:nvPicPr>
            <pic:cNvPr id="{shape_id}" name="{name}"/>
            <pic:cNvPicPr>
              <a:picLocks noChangeAspect="1" noChangeArrowheads="1"/>
            </pic:cNvPicPr>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip r:embed="{rid}"/>
            <a:stretch>
              <a:fillRect/>
            </a:stretch>
          </pic:blipFill>
          <pic:spPr bwMode="auto">
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
            <a:noFill/>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # First paragraph (anchor for floating shapes)
    para0 = doc.add_paragraph('')

    # --- Step 1: Add rectangle shape as floating anchor ---
    rect_x = cm_to_emu(0.2)    # Not at 0,0 - agent must reposition
    rect_y = cm_to_emu(0.3)
    rect_w = cm_to_emu(17.0)
    rect_h = cm_to_emu(2.5)

    rect_xml = build_rect_anchor_xml(rect_x, rect_y, rect_w, rect_h,
                                     shape_id=1, name='Rectangle1',
                                     fill_hex='1F3864')
    rect_elem = etree.fromstring(rect_xml)
    run1 = OxmlElement('w:r')
    run1.append(rect_elem)
    para0._p.append(run1)

    # --- Step 2: Add logo image as floating anchor ---
    # First add inline to get rId, then replace with anchor XML
    logo_bytes = make_png_bytes(60, 60, 0, 120, 180)
    img_stream = io.BytesIO(logo_bytes)

    # Add inline picture to get rId registered
    tmp_run = para0.add_run()
    tmp_run.add_picture(img_stream, width=Cm(2), height=Cm(2))

    # Find the rId from the inline drawing that was just added
    pic_xml_str = para0._p.xml
    match = re.search(r'r:embed="(rId\d+)"', pic_xml_str)
    if not match:
        raise ValueError('Could not find rId for logo image')
    logo_rid = match.group(1)

    # Remove the inline drawing run (last run in para0)
    inline_runs = para0._p.findall(qn('w:r'))
    para0._p.remove(inline_runs[-1])

    # Now add the logo as a floating anchor
    logo_x = cm_to_emu(0.3)    # Not at 0,0 - separate from rectangle anchor
    logo_y = cm_to_emu(0.4)
    logo_w = cm_to_emu(2.0)
    logo_h = cm_to_emu(2.0)

    img_xml = build_image_anchor_xml(logo_x, logo_y, logo_w, logo_h,
                                      shape_id=2, name='Logo1', rid=logo_rid)
    img_elem = etree.fromstring(img_xml)
    run2 = OxmlElement('w:r')
    run2.append(img_elem)
    para0._p.append(run2)

    # --- Document content ---
    doc.add_paragraph('')

    h1 = doc.add_heading('', level=1)
    run_h1 = h1.add_run('Quarterly Business Review')
    run_h1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph(
        'This report summarizes the key performance indicators and financial results '
        'for Q4 2025. Our team has achieved significant milestones across all '
        'business units, demonstrating strong growth and operational excellence.'
    )

    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'Total revenue for the quarter reached $4.2 million, representing a 18% '
        'year-over-year increase. Operating margins improved to 24.5%, up from '
        '21.2% in the same period last year.'
    )

    doc.add_heading('Department Performance', level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Department'
    hdr_cells[1].text = 'Q4 Revenue'
    hdr_cells[2].text = 'YoY Growth'
    hdr_cells[3].text = 'Headcount'

    departments = [
        ('Engineering', '$1,850,000', '+22%', '47'),
        ('Sales & Marketing', '$1,420,000', '+15%', '31'),
        ('Customer Success', '$620,000', '+18%', '22'),
        ('Operations', '$310,000', '+12%', '18'),
    ]
    for dept, revenue, growth, headcount in departments:
        row_cells = table.add_row().cells
        row_cells[0].text = dept
        row_cells[1].text = revenue
        row_cells[2].text = growth
        row_cells[3].text = headcount

    doc.add_paragraph('')
    doc.add_heading('Key Highlights', level=2)
    doc.add_paragraph(
        '• Product launch in APAC market exceeded projections by 32%\n'
        '• Customer retention rate improved to 94.7% from 91.2%\n'
        '• New enterprise contracts signed: 14 (avg. contract value $89,000)\n'
        '• Employee satisfaction score: 4.3/5.0 (up from 4.1 last quarter)'
    )

    doc.add_heading('Strategic Initiatives for Q1 2026', level=2)
    doc.add_paragraph(
        'The leadership team has approved three major strategic initiatives for '
        'the upcoming quarter: expansion into the European market, launch of our '
        'AI-powered analytics platform, and restructuring of the partner program '
        'to improve channel revenue.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
