"""
Reward Script: Insert section break before 'Appendix A' with restarted page numbering and custom footer
Task ID: writer_biz_047
Domain: libreoffice_writer
Scoring:
  Component 1: Document has 2+ sections (section break inserted) — 0.25 points
  Component 2: Appendix section restarts page numbering at 1 — 0.25 points
  Component 3: Appendix section footer contains 'Appendix - Page' text — 0.30 points
  Component 4: Appendix section footer has PAGE field code — 0.20 points
"""

import os
import re
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_047'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = doc.sections
    num_sections = len(sections)

    # Component 1: Document has 2+ sections (section break was inserted) (0.25 points)
    # Initial has 1 section; golden should have 2+.
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (section break inserted) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected 2+ sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Appendix section (last section) restarts page numbering at 1 (0.25 points)
    # Initial has no pgNumType restart; golden has pgNumType start=1 on the appendix section.
    try:
        if num_sections >= 2:
            appendix_section = sections[-1]
            sectPr = appendix_section._sectPr
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                start_val = pgNumType.get(qn('w:start'))
                if start_val == '1':
                    print(f"PASS: Component 2 — Appendix section pgNumType start=1 (page numbering restarted) (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 2 — pgNumType start={start_val}, expected '1'")
            else:
                print(f"FAIL: Component 2 — No pgNumType element found in appendix section")
        else:
            print(f"FAIL: Component 2 — Cannot check page numbering without 2+ sections")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Appendix section footer contains 'Appendix - Page' text (0.30 points)
    # Initial footer only says 'Page '; golden appendix footer says 'Appendix - Page '.
    try:
        if num_sections >= 2:
            appendix_section = sections[-1]
            footer = appendix_section.footer
            footer_text = ''
            if footer.paragraphs:
                footer_text = ' '.join(fp.text for fp in footer.paragraphs)

            # Check for the pattern "Appendix" and "Page" in footer text
            if 'Appendix' in footer_text and 'Page' in footer_text:
                # More specifically check for "Appendix - Page" pattern
                if re.search(r'Appendix\s*[-\u2013\u2014]\s*Page', footer_text):
                    print(f"PASS: Component 3 — Footer text contains 'Appendix - Page' pattern: {repr(footer_text)} (0.30 pts)")
                    total_score += 0.30
                else:
                    print(f"PARTIAL: Component 3 — Footer has 'Appendix' and 'Page' but not in expected format: {repr(footer_text)} (0.15 pts)")
                    total_score += 0.15
            else:
                print(f"FAIL: Component 3 — Footer text {repr(footer_text)} does not contain 'Appendix' and 'Page'")
        else:
            print(f"FAIL: Component 3 — Cannot check footer without 2+ sections")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Appendix section footer has PAGE field code (0.20 points)
    # This ensures the footer actually shows dynamic page numbers, not just static text.
    try:
        if num_sections >= 2:
            appendix_section = sections[-1]
            footer = appendix_section.footer
            has_page_field = False
            if footer.paragraphs:
                for fp in footer.paragraphs:
                    xml_str = fp._element.xml
                    if 'instrText' in xml_str:
                        instrs = re.findall(r'<w:instrText[^>]*>([^<]+)</w:instrText>', xml_str)
                        for instr in instrs:
                            if 'PAGE' in instr.upper():
                                has_page_field = True
                                break
                    if has_page_field:
                        break

            if has_page_field:
                print(f"PASS: Component 4 — Appendix footer contains PAGE field code (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — No PAGE field code found in appendix footer")
        else:
            print(f"FAIL: Component 4 — Cannot check footer field without 2+ sections")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
