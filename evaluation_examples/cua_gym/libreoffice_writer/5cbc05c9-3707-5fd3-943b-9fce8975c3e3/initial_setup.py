"""
Initial Setup: Insert a section break before 'Appendix A' and restart page numbering
Task ID: writer_biz_047
Domain: libreoffice_writer

Creates a multi-page business document with continuous page numbering.
The document has main body content (pages 1-10) followed by Appendix A content.
All pages use Default Page Style with continuous numbering and "Page X" footer.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section, prefix="Page "):
    """Add a footer with page number field code."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Clear existing content
    for run in fp.runs:
        run.clear()

    # Add prefix text
    run_prefix = fp.add_run(prefix)
    run_prefix.font.size = Pt(10)
    run_prefix.font.name = "Calibri"

    # Add PAGE field code
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    instr.set(qn('xml:space'), 'preserve')
    r2._element.append(instr)

    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Set default page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Add footer with "Page X" to the default section
    add_page_number_footer(section, prefix="Page ")

    # ===== TITLE PAGE =====
    title = doc.add_heading("Meridian Technologies Annual Business Review", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Fiscal Year 2025 - Strategic Planning Report")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run("Prepared by: Corporate Strategy Division\nMarch 15, 2025")
    run.font.size = Pt(11)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS PAGE =====
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary ......................................... 3",
        "2. Market Analysis ............................................ 4",
        "3. Financial Performance ...................................... 5",
        "4. Product Development Pipeline ............................... 6",
        "5. Human Resources Overview ................................... 7",
        "6. Technology Infrastructure .................................. 8",
        "7. Risk Assessment ............................................ 9",
        "8. Strategic Recommendations .................................. 10",
        "Appendix A: Supplementary Data Tables ......................... 11",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ===== SECTION 1: EXECUTIVE SUMMARY =====
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Meridian Technologies concluded fiscal year 2025 with record-breaking revenue of $487.3 million, "
        "representing a 23.4% year-over-year increase. This growth was primarily driven by our expanded "
        "cloud services portfolio and strategic acquisitions in the cybersecurity sector. The company "
        "successfully onboarded 1,847 new enterprise clients while maintaining a 94.2% retention rate "
        "among existing customers."
    )
    doc.add_paragraph(
        "Operating margins improved to 18.7%, up from 15.3% in the prior year, reflecting our ongoing "
        "commitment to operational efficiency and automation initiatives. Research and development "
        "expenditure grew to $62.8 million, representing 12.9% of total revenue, as we accelerated "
        "investment in artificial intelligence and machine learning capabilities."
    )
    doc.add_paragraph(
        "Our workforce expanded to 3,240 full-time employees across 12 global offices, with notable "
        "growth in our Singapore and Berlin engineering centers. Employee satisfaction scores reached "
        "an all-time high of 4.3 out of 5.0, driven by enhanced benefits packages and flexible work "
        "arrangements implemented during the fiscal year."
    )

    doc.add_page_break()

    # ===== SECTION 2: MARKET ANALYSIS =====
    doc.add_heading("2. Market Analysis", level=1)
    doc.add_paragraph(
        "The global enterprise software market reached $672 billion in 2025, with cloud-native solutions "
        "accounting for 41% of total spending. Meridian's addressable market expanded significantly "
        "following our entry into the healthcare IT and financial services compliance verticals."
    )
    doc.add_heading("2.1 Competitive Landscape", level=2)
    doc.add_paragraph(
        "Key competitors include Nextera Solutions, Pinnacle Systems, and Vertex Digital. Meridian "
        "maintained its position as the third-largest provider in the North American market with "
        "8.3% market share, up from 6.7% in the prior year. Our Net Promoter Score of 72 continues "
        "to lead the industry average of 54."
    )
    doc.add_heading("2.2 Regional Performance", level=2)
    doc.add_paragraph(
        "North America contributed $312.5 million (64.1%) of total revenue, followed by Europe at "
        "$109.8 million (22.5%), and Asia-Pacific at $65.0 million (13.3%). The Asia-Pacific region "
        "showed the strongest growth at 38.2%, driven by expansion in Japan and Australia."
    )

    doc.add_page_break()

    # ===== SECTION 3: FINANCIAL PERFORMANCE =====
    doc.add_heading("3. Financial Performance", level=1)
    doc.add_paragraph(
        "Revenue for fiscal year 2025 totaled $487.3 million, with recurring subscription revenue "
        "accounting for $341.1 million (70.0%). Professional services contributed $97.5 million "
        "(20.0%), and perpetual licensing accounted for $48.7 million (10.0%)."
    )

    # Financial table
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Metric", "FY 2024", "FY 2025", "Change"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    fin_data = [
        ["Total Revenue", "$394.9M", "$487.3M", "+23.4%"],
        ["Operating Income", "$60.4M", "$91.1M", "+50.8%"],
        ["Net Income", "$47.2M", "$73.4M", "+55.5%"],
        ["Free Cash Flow", "$52.8M", "$84.6M", "+60.2%"],
        ["Earnings Per Share", "$2.14", "$3.28", "+53.3%"],
    ]
    for r, row_data in enumerate(fin_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph("")  # spacer

    doc.add_page_break()

    # ===== SECTION 4: PRODUCT DEVELOPMENT PIPELINE =====
    doc.add_heading("4. Product Development Pipeline", level=1)
    doc.add_paragraph(
        "Our product roadmap for 2026 encompasses 14 major feature releases across three core "
        "platforms: Meridian Cloud Suite, Meridian Security Shield, and Meridian Analytics Pro. "
        "Investments in AI-driven automation represent the single largest development initiative."
    )
    doc.add_heading("4.1 Cloud Suite Enhancements", level=2)
    doc.add_paragraph(
        "Planned improvements include multi-region failover capabilities, enhanced API gateway "
        "performance (targeting 99.999% uptime), and integration with 45 additional third-party "
        "services. The new microservices architecture is expected to reduce deployment times by 60%."
    )
    doc.add_heading("4.2 Security Shield Updates", level=2)
    doc.add_paragraph(
        "Version 4.0 of Security Shield will introduce behavioral threat detection powered by "
        "proprietary machine learning models trained on 2.3 billion security events. Zero-trust "
        "network architecture support and automated incident response playbooks are also planned "
        "for Q3 2026 release."
    )

    doc.add_page_break()

    # ===== SECTION 5: HUMAN RESOURCES OVERVIEW =====
    doc.add_heading("5. Human Resources Overview", level=1)
    doc.add_paragraph(
        "Total headcount grew from 2,810 to 3,240 employees, a net increase of 430 positions. "
        "Engineering remained the largest department with 1,458 employees (45%), followed by "
        "Sales and Marketing at 778 (24%), Customer Success at 486 (15%), and Corporate Functions "
        "at 518 (16%)."
    )
    doc.add_paragraph(
        "Voluntary turnover decreased to 11.2% from 14.8%, well below the technology industry "
        "average of 18.5%. Key retention initiatives included expanded equity compensation programs, "
        "a four-day work week pilot in select offices, and enhanced parental leave policies."
    )

    doc.add_page_break()

    # ===== SECTION 6: TECHNOLOGY INFRASTRUCTURE =====
    doc.add_heading("6. Technology Infrastructure", level=1)
    doc.add_paragraph(
        "Infrastructure spending totaled $38.4 million in FY 2025, representing 7.9% of revenue. "
        "Major investments included migration of 78% of workloads to multi-cloud environments "
        "(AWS, Azure, and GCP), deployment of new edge computing nodes in 8 metropolitan areas, "
        "and implementation of a company-wide zero-trust security framework."
    )
    doc.add_paragraph(
        "System reliability improved significantly with platform uptime reaching 99.97%, up from "
        "99.91% in the prior year. Mean time to resolution for critical incidents decreased from "
        "47 minutes to 23 minutes following the deployment of AI-powered monitoring and automated "
        "remediation tools."
    )

    doc.add_page_break()

    # ===== SECTION 7: RISK ASSESSMENT =====
    doc.add_heading("7. Risk Assessment", level=1)
    doc.add_paragraph(
        "The Enterprise Risk Management team identified 23 material risks across four categories: "
        "operational, financial, strategic, and compliance. The top five risks by potential impact "
        "are summarized below."
    )
    risk_items = [
        "Cybersecurity threats and data breach potential (Critical)",
        "Regulatory compliance in new market verticals (High)",
        "Key personnel retention in competitive labor market (High)",
        "Supply chain disruptions affecting hardware procurement (Medium)",
        "Foreign exchange exposure from international operations (Medium)",
    ]
    for item in risk_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Mitigation strategies have been developed for each identified risk, with quarterly review "
        "cycles established for the top ten risks. Insurance coverage was expanded to include "
        "comprehensive cyber liability and business interruption policies."
    )

    doc.add_page_break()

    # ===== SECTION 8: STRATEGIC RECOMMENDATIONS =====
    doc.add_heading("8. Strategic Recommendations", level=1)
    doc.add_paragraph(
        "Based on the analysis presented in this report, the Strategic Planning Division recommends "
        "the following priorities for fiscal year 2026:"
    )
    recommendations = [
        "Accelerate AI integration across all product lines with a dedicated $25M investment fund",
        "Expand Asia-Pacific presence with new offices in Seoul and Mumbai by Q2 2026",
        "Pursue strategic acquisition of a mid-market analytics company valued at $50-80M",
        "Launch an enterprise developer program targeting 10,000 registered developers by year-end",
        "Implement sustainability initiatives to achieve carbon-neutral operations by 2027",
    ]
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f"{i}. {rec}")

    doc.add_paragraph(
        "These recommendations align with our five-year strategic vision and are projected to "
        "drive revenue growth of 25-30% in fiscal year 2026 while maintaining operating margins "
        "above 18%."
    )

    doc.add_page_break()

    # ===== APPENDIX A =====
    doc.add_heading("Appendix A", level=1)
    doc.add_heading("Supplementary Data Tables", level=2)

    doc.add_paragraph(
        "The following tables provide detailed breakdowns of key metrics referenced in the main "
        "body of this report."
    )

    # Appendix Table 1: Quarterly Revenue
    doc.add_heading("A.1 Quarterly Revenue Breakdown", level=3)
    tbl1 = doc.add_table(rows=5, cols=5)
    tbl1.style = "Table Grid"
    q_headers = ["Quarter", "Cloud Services", "Professional Services", "Licensing", "Total"]
    for i, h in enumerate(q_headers):
        cell = tbl1.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    q_data = [
        ["Q1 2025", "$78.2M", "$22.1M", "$11.4M", "$111.7M"],
        ["Q2 2025", "$82.5M", "$23.8M", "$12.1M", "$118.4M"],
        ["Q3 2025", "$88.9M", "$25.2M", "$12.7M", "$126.8M"],
        ["Q4 2025", "$91.5M", "$26.4M", "$12.5M", "$130.4M"],
    ]
    for r, row_data in enumerate(q_data, 1):
        for c, val in enumerate(row_data):
            tbl1.cell(r, c).text = val

    doc.add_paragraph("")  # spacer

    # Appendix Table 2: Regional Headcount
    doc.add_heading("A.2 Regional Headcount Distribution", level=3)
    tbl2 = doc.add_table(rows=6, cols=4)
    tbl2.style = "Table Grid"
    rh_headers = ["Region", "FY 2024", "FY 2025", "Growth"]
    for i, h in enumerate(rh_headers):
        cell = tbl2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rh_data = [
        ["North America", "1,580", "1,782", "+12.8%"],
        ["Europe", "720", "842", "+16.9%"],
        ["Asia-Pacific", "380", "456", "+20.0%"],
        ["Latin America", "90", "112", "+24.4%"],
        ["Middle East & Africa", "40", "48", "+20.0%"],
    ]
    for r, row_data in enumerate(rh_data, 1):
        for c, val in enumerate(row_data):
            tbl2.cell(r, c).text = val

    doc.add_paragraph("")  # spacer

    # Appendix Table 3: Product Performance
    doc.add_heading("A.3 Product Line Performance", level=3)
    tbl3 = doc.add_table(rows=5, cols=4)
    tbl3.style = "Table Grid"
    pp_headers = ["Product", "Revenue", "Clients", "Satisfaction"]
    for i, h in enumerate(pp_headers):
        cell = tbl3.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    pp_data = [
        ["Cloud Suite", "$245.8M", "2,134", "4.4/5.0"],
        ["Security Shield", "$128.7M", "1,567", "4.2/5.0"],
        ["Analytics Pro", "$78.3M", "892", "4.1/5.0"],
        ["Legacy Products", "$34.5M", "423", "3.8/5.0"],
    ]
    for r, row_data in enumerate(pp_data, 1):
        for c, val in enumerate(row_data):
            tbl3.cell(r, c).text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
