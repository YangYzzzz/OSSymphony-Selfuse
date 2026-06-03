"""
Initial Setup: Create a Writer document with a thesis outline as plain paragraphs
Task ID: writer_acad_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_072'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Thesis Outline: The Impact of Artificial Intelligence on Modern Healthcare Systems', level=0)

    # Brief intro paragraph
    intro = doc.add_paragraph(
        'This thesis examines how artificial intelligence technologies are transforming '
        'healthcare delivery, diagnostics, and patient outcomes across developed nations. '
        'The following outline presents the structure of the research.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # All outline items as plain paragraphs - NO numbering, NO indentation, NO list styles
    # The task is for the agent to convert these into a two-level numbered list

    outline_items = [
        # Level 1 topics and their Level 2 subtopics - all as plain paragraphs
        ("Introduction", True),
        ("Background and Motivation", False),
        ("Research Questions and Objectives", False),
        ("Scope and Limitations of the Study", False),

        ("Literature Review", True),
        ("Historical Development of AI in Medicine", False),
        ("Current Applications of Machine Learning in Diagnostics", False),
        ("Ethical Frameworks for AI-Assisted Decision Making", False),

        ("Methodology", True),
        ("Research Design and Approach", False),
        ("Data Collection Methods and Sources", False),
        ("Analytical Framework and Statistical Tools", False),

        ("Results and Analysis", True),
        ("Quantitative Findings from Hospital Case Studies", False),
        ("Qualitative Assessment of Clinician Perspectives", False),

        ("Discussion", True),
        ("Interpretation of Key Findings", False),
        ("Comparison with Existing Literature", False),
        ("Implications for Healthcare Policy", False),

        ("Conclusion and Future Directions", True),
        ("Summary of Contributions", False),
        ("Recommendations for Practitioners", False),
        ("Proposed Areas for Further Research", False),
    ]

    for text, is_main in outline_items:
        para = doc.add_paragraph(text)
        # Main topics get slightly larger font but NO numbering
        if is_main:
            for run in para.runs:
                run.font.size = Pt(13)
                run.font.bold = True
        else:
            for run in para.runs:
                run.font.size = Pt(11)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
