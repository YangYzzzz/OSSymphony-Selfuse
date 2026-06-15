"""
Initial Setup: Delete the third column from the table and add a new header row at the top.
Task ID: osworld_writer_table_editing_002
Domain: libreoffice_writer

Creates a document with a 5-row, 4-column data table (no header row).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_table_editing_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    doc.add_heading('Employee Records — Q1 2025', level=1)
    doc.add_paragraph(
        'The following table lists department employees with their performance metrics for the first quarter.'
    )

    # 5-row, 4-column data table — NO header row
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Realistic employee data: Name | Department | Review Score | Start Date
    data = [
        ['Sarah Chen',      'Engineering',  '4.7',  '2021-03-15'],
        ['Marcus Johnson',  'Marketing',    '4.2',  '2020-06-01'],
        ['Priya Patel',     'Finance',      '4.5',  '2019-11-20'],
        ['Tom Reyes',       'Operations',   '3.9',  '2022-02-08'],
        ['Linda Okafor',    'Engineering',  '4.8',  '2018-07-30'],
    ]

    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = value

    doc.add_paragraph('')
    doc.add_paragraph('Data last updated: March 2025.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
