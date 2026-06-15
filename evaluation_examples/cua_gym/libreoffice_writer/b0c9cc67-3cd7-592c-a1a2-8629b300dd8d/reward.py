"""
Reward Script: Delete the third column from the table in this document and
               then add a new row at the top to serve as a header row.
Task ID: osworld_writer_table_editing_002
Domain: libreoffice_writer
Scoring:
  Component 1: Table has exactly 3 columns (third column removed)       — 0.4 pts
  Component 2: Table has exactly 6 rows (new header row added at top)   — 0.3 pts
  Component 3: Data rows retain correct content (original cols 1,2,4)   — 0.3 pts
Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_table_editing_002'

# Expected data rows after task completion (original columns 1, 2, 4 — i.e., Name, Dept, Date)
EXPECTED_DATA_ROWS = [
    ('Sarah Chen',     'Engineering', '2021-03-15'),
    ('Marcus Johnson', 'Marketing',   '2020-06-01'),
    ('Priya Patel',    'Finance',     '2019-11-20'),
    ('Tom Reyes',      'Operations',  '2022-02-08'),
    ('Linda Okafor',   'Engineering', '2018-07-30'),
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Initial state: 1 table, 5 rows x 4 columns (Name, Department, Rating, Start Date)
    Expected state: 1 table, 6 rows x 3 columns (empty header row + 5 data rows, Rating column gone)
    """
    total_score = 0.0

    # Precondition gate: file must load and have exactly one table
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    print(f"INFO: Table dimensions: {num_rows} rows x {num_cols} cols")

    # Component 1: Table has exactly 3 columns — third column (Rating) was deleted (0.4 pts)
    try:
        if num_cols == 3:
            print(f"PASS: Component 1 — Table has exactly 3 columns (Rating column deleted) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected 3 columns, found {num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table has exactly 6 rows — new header row was added at top (0.3 pts)
    try:
        if num_rows == 6:
            print(f"PASS: Component 2 — Table has exactly 6 rows (header row added) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected 6 rows, found {num_rows}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Row 0 is the empty header row AND data rows 1-5 retain correct content (0.3 pts)
    # This verifies both the header row structure and data integrity
    try:
        if num_rows >= 6 and num_cols == 3:
            # Check row 0 is the (empty) header row
            header_row_cells = [table.cell(0, c).text.strip() for c in range(3)]
            # The header row should be newly added (empty or contain headers — per task context it's empty)
            # We accept empty header row as correct
            header_is_empty = all(c == '' for c in header_row_cells)

            # Check data rows 1..5 match expected (without the old 3rd/Rating column)
            data_matches = 0
            for row_idx, (expected_name, expected_dept, expected_date) in enumerate(EXPECTED_DATA_ROWS, start=1):
                actual_name = table.cell(row_idx, 0).text.strip()
                actual_dept = table.cell(row_idx, 1).text.strip()
                actual_date = table.cell(row_idx, 2).text.strip()
                if (actual_name == expected_name and
                        actual_dept == expected_dept and
                        actual_date == expected_date):
                    data_matches += 1
                else:
                    print(f"  MISMATCH row {row_idx}: got ({actual_name!r}, {actual_dept!r}, {actual_date!r}), "
                          f"expected ({expected_name!r}, {expected_dept!r}, {expected_date!r})")

            all_data_correct = (data_matches == 5)

            if header_is_empty and all_data_correct:
                print(f"PASS: Component 3 — Empty header row at row 0 and all 5 data rows correct (0.3 pts)")
                total_score += 0.3
            elif all_data_correct and not header_is_empty:
                print(f"PARTIAL-FAIL: Component 3 — Data rows correct but header row is not empty: {header_row_cells}")
                # Still award partial credit for data being correct (but we set exact 0 here since it's a single component)
                print(f"FAIL: Component 3 — Header row content mismatch. Expected empty row, got {header_row_cells}")
            elif header_is_empty and not all_data_correct:
                print(f"FAIL: Component 3 — Header row is empty but only {data_matches}/5 data rows match expected content")
            else:
                print(f"FAIL: Component 3 — Header row not empty ({header_row_cells}) AND {data_matches}/5 data rows correct")
        else:
            print(f"SKIP: Component 3 — Skipped because table dimensions are incorrect ({num_rows} rows x {num_cols} cols)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
