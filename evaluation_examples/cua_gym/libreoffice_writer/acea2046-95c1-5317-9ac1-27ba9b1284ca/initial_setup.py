"""
Initial Setup: Technical proposal document with key terms for indexing
Task ID: writer_biz_070
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_070'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Enterprise Cloud Migration Technical Proposal', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared for Meridian Financial Services\nMarch 2025')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()  # spacer

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This proposal outlines a comprehensive strategy for migrating Meridian Financial '
        'Services from its legacy on-premises infrastructure to a modern Cloud Infrastructure '
        'platform. The migration encompasses all critical business systems, customer-facing '
        'applications, and internal operational tools.'
    )
    doc.add_paragraph(
        'Our approach prioritizes minimal downtime through a phased Data Migration strategy '
        'that ensures business continuity while systematically transitioning workloads. Each '
        'phase includes rigorous testing, validation, and rollback procedures to protect '
        'against data loss or service interruption.'
    )
    doc.add_paragraph(
        'The proposed solution includes enterprise-grade SLA commitments backed by '
        'contractual guarantees, ensuring 99.95% uptime for all Tier-1 services and 99.9% '
        'for Tier-2 services throughout the migration period and beyond.'
    )

    # --- Section 2: Technical Architecture ---
    doc.add_heading('2. Technical Architecture', level=1)

    doc.add_heading('2.1 Cloud Infrastructure Design', level=2)
    doc.add_paragraph(
        'The target Cloud Infrastructure architecture utilizes a multi-region deployment '
        'model with automatic failover capabilities. Primary workloads will be hosted in '
        'the US-East region with secondary replicas in US-West, providing geographic '
        'redundancy for all critical services.'
    )
    doc.add_paragraph(
        'Network topology follows a hub-and-spoke design pattern with dedicated virtual '
        'private clouds for each business unit. Inter-VPC communication is secured through '
        'encrypted transit gateways, and all external traffic routes through web application '
        'firewalls and DDoS mitigation layers.'
    )

    doc.add_heading('2.2 API Integration Framework', level=2)
    doc.add_paragraph(
        'Seamless connectivity between legacy and cloud-native systems will be achieved '
        'through a centralized API Integration gateway. This middleware layer handles '
        'protocol translation, rate limiting, authentication, and request routing across '
        'all connected services.'
    )
    doc.add_paragraph(
        'The API Integration platform supports RESTful, GraphQL, and gRPC protocols, '
        'enabling gradual migration of individual services without disrupting dependent '
        'applications. A comprehensive API catalog will document all available endpoints, '
        'authentication requirements, and usage quotas.'
    )

    # --- Section 3: Data Migration Strategy ---
    doc.add_heading('3. Data Migration Strategy', level=1)
    doc.add_paragraph(
        'The Data Migration process follows a three-phase approach designed to minimize '
        'risk and ensure data integrity throughout the transition. Phase 1 addresses '
        'non-critical archival data, Phase 2 migrates operational databases, and Phase 3 '
        'handles real-time transactional systems.'
    )

    # Table: Migration Phases
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['Phase', 'Data Category', 'Volume (TB)', 'Timeline']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    phase_data = [
        ['Phase 1', 'Archival & Historical Records', '45', 'Weeks 1-4'],
        ['Phase 2', 'Operational Databases & Analytics', '120', 'Weeks 5-10'],
        ['Phase 3', 'Real-time Transactional Systems', '85', 'Weeks 11-14'],
    ]
    for r, row_data in enumerate(phase_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'Each Data Migration phase includes automated validation scripts that compare '
        'source and target checksums, row counts, and referential integrity constraints. '
        'Any discrepancies trigger automatic rollback procedures and alert the migration '
        'operations team.'
    )

    # --- Section 4: Service Level Agreements ---
    doc.add_heading('4. Service Level Agreements', level=1)
    doc.add_paragraph(
        'All services covered under this proposal are governed by a comprehensive SLA '
        'framework. The SLA defines measurable performance targets, escalation procedures, '
        'and financial remedies for any service disruptions that exceed defined thresholds.'
    )

    # Table: SLA Tiers
    sla_table = doc.add_table(rows=4, cols=3)
    sla_table.style = 'Table Grid'
    sla_headers = ['SLA Tier', 'Uptime Guarantee', 'Response Time']
    for i, h in enumerate(sla_headers):
        cell = sla_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    sla_data = [
        ['Tier 1 - Critical', '99.95%', '15 minutes'],
        ['Tier 2 - Standard', '99.9%', '1 hour'],
        ['Tier 3 - Non-Critical', '99.5%', '4 hours'],
    ]
    for r, row_data in enumerate(sla_data, 1):
        for c, val in enumerate(row_data):
            sla_table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'Monthly SLA compliance reports will be generated automatically and shared with '
        'Meridian stakeholders. Any month where SLA targets are not met will trigger '
        'service credits as outlined in Appendix B of the master services agreement.'
    )

    # --- Section 5: Disaster Recovery ---
    doc.add_heading('5. Disaster Recovery & Business Continuity', level=1)
    doc.add_paragraph(
        'A robust Disaster Recovery plan is integral to the proposed Cloud Infrastructure. '
        'The Disaster Recovery strategy employs active-passive replication with automated '
        'failover, ensuring recovery point objectives (RPO) of less than 15 minutes and '
        'recovery time objectives (RTO) of less than 1 hour for all Tier-1 services.'
    )
    doc.add_paragraph(
        'Quarterly Disaster Recovery drills will be conducted to validate failover '
        'procedures and measure actual recovery times against defined SLA targets. Results '
        'from each drill are documented and used to refine the Disaster Recovery playbook.'
    )
    doc.add_paragraph(
        'The backup architecture leverages incremental snapshots stored across multiple '
        'geographic regions, with full backups performed weekly. Backup integrity is '
        'verified through automated restoration testing on isolated Cloud Infrastructure '
        'instances.'
    )

    # --- Section 6: Implementation Timeline ---
    doc.add_heading('6. Implementation Timeline', level=1)
    doc.add_paragraph(
        'The complete migration is projected to span 18 weeks, with key milestones tied '
        'to each Data Migration phase. API Integration work runs in parallel, beginning '
        'in Week 2 and continuing through final cutover.'
    )

    # Timeline table
    timeline_table = doc.add_table(rows=6, cols=3)
    timeline_table.style = 'Table Grid'
    tl_headers = ['Milestone', 'Timeline', 'Dependencies']
    for i, h in enumerate(tl_headers):
        cell = timeline_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    tl_data = [
        ['Cloud Infrastructure provisioning', 'Weeks 1-2', 'Network design approval'],
        ['API Integration gateway deployment', 'Weeks 2-4', 'Cloud Infrastructure ready'],
        ['Data Migration Phase 1', 'Weeks 1-4', 'Storage provisioning'],
        ['Data Migration Phase 2-3', 'Weeks 5-14', 'Phase 1 validation'],
        ['Disaster Recovery testing', 'Weeks 15-16', 'All migrations complete'],
    ]
    for r, row_data in enumerate(tl_data, 1):
        for c, val in enumerate(row_data):
            timeline_table.cell(r, c).text = val

    # --- Section 7: Budget & Resources ---
    doc.add_heading('7. Budget & Resources', level=1)
    doc.add_paragraph(
        'The total project budget is estimated at $2.4 million, allocated across '
        'Cloud Infrastructure costs ($890,000), API Integration development ($520,000), '
        'Data Migration services ($460,000), Disaster Recovery setup ($310,000), and '
        'project management and SLA monitoring tools ($220,000).'
    )

    # Budget table
    budget_table = doc.add_table(rows=6, cols=2)
    budget_table.style = 'Table Grid'
    b_headers = ['Category', 'Budget']
    for i, h in enumerate(b_headers):
        cell = budget_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    b_data = [
        ['Cloud Infrastructure', '$890,000'],
        ['API Integration', '$520,000'],
        ['Data Migration', '$460,000'],
        ['Disaster Recovery', '$310,000'],
        ['SLA Monitoring & PM', '$220,000'],
    ]
    for r, row_data in enumerate(b_data, 1):
        for c, val in enumerate(row_data):
            budget_table.cell(r, c).text = val

    doc.add_paragraph()

    # --- Section 8: Conclusion ---
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph(
        'This technical proposal presents a comprehensive plan for Meridian Financial '
        'Services to modernize its technology infrastructure through a carefully managed '
        'cloud migration. Our commitment to stringent SLA standards, thorough Data Migration '
        'practices, reliable Disaster Recovery capabilities, and flexible API Integration '
        'architecture ensures a smooth transition with minimal business disruption.'
    )
    doc.add_paragraph(
        'We look forward to partnering with Meridian Financial Services on this '
        'transformative initiative and remain available to discuss any aspects of this '
        'proposal in further detail.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
