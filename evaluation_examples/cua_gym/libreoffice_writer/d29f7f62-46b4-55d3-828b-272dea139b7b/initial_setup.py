"""
Initial Setup: Create a Writer document with Heading 1 and Heading 2 styles applied,
but Heading 2 has no outline numbering.
Task ID: writer_tech_089
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_089'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Cloud Infrastructure Deployment Guide', level=0)

    # Introduction paragraph
    doc.add_paragraph(
        'This document provides comprehensive guidelines for deploying and managing '
        'cloud infrastructure components across multiple environments. It covers '
        'initial setup, configuration management, monitoring, and incident response '
        'procedures for the engineering team.'
    )

    # Chapter 1
    doc.add_heading('Environment Setup', level=1)

    doc.add_paragraph(
        'Before deploying any services, the target environment must be properly '
        'configured. This includes provisioning virtual machines, setting up network '
        'security groups, and configuring DNS records for service discovery.'
    )

    doc.add_heading('Prerequisites', level=2)

    doc.add_paragraph(
        'Ensure the following prerequisites are met before proceeding with the '
        'environment setup:'
    )
    doc.add_paragraph('AWS CLI version 2.x or later installed and configured', style='List Bullet')
    doc.add_paragraph('Terraform v1.5+ with the AWS provider plugin', style='List Bullet')
    doc.add_paragraph('Docker Engine 24.0 or later on all build machines', style='List Bullet')
    doc.add_paragraph('Valid SSL certificates for all public-facing endpoints', style='List Bullet')

    doc.add_heading('Network Configuration', level=2)

    doc.add_paragraph(
        'The production VPC uses a /16 CIDR block (10.0.0.0/16) with three availability '
        'zones. Each availability zone contains a public subnet (/24) for load balancers '
        'and a private subnet (/22) for application servers and databases.'
    )

    doc.add_paragraph(
        'Security groups must be configured to allow inbound traffic on port 443 (HTTPS) '
        'from the public internet, and port 8080 for internal service-to-service '
        'communication within the VPC.'
    )

    doc.add_heading('IAM Roles and Policies', level=2)

    doc.add_paragraph(
        'Each microservice runs with a dedicated IAM role following the principle of '
        'least privilege. The deployment pipeline uses a cross-account assume-role '
        'pattern to access resources in staging and production accounts.'
    )

    # Chapter 2
    doc.add_heading('Service Deployment', level=1)

    doc.add_paragraph(
        'Services are deployed using a blue-green deployment strategy to minimize '
        'downtime and enable instant rollback. The CI/CD pipeline handles building '
        'Docker images, running integration tests, and promoting artifacts through '
        'the deployment stages.'
    )

    doc.add_heading('Container Registry Setup', level=2)

    doc.add_paragraph(
        'All Docker images are stored in Amazon ECR with lifecycle policies that '
        'automatically clean up untagged images older than 30 days. Each service '
        'has its own repository following the naming convention: '
        'company-name/service-name.'
    )

    doc.add_heading('Kubernetes Configuration', level=2)

    doc.add_paragraph(
        'The EKS cluster runs Kubernetes 1.28 with managed node groups. Pod resource '
        'requests and limits are enforced via LimitRange objects in each namespace. '
        'Horizontal Pod Autoscaler (HPA) is configured for all stateless services '
        'with target CPU utilization of 70%.'
    )

    # Chapter 3
    doc.add_heading('Monitoring and Alerting', level=1)

    doc.add_paragraph(
        'Comprehensive monitoring is essential for maintaining service reliability. '
        'The monitoring stack includes Prometheus for metrics collection, Grafana '
        'for visualization, and PagerDuty for incident alerting.'
    )

    doc.add_heading('Metrics Collection', level=2)

    doc.add_paragraph(
        'Prometheus scrapes metrics from all service endpoints at 15-second intervals. '
        'Custom metrics are exposed via the /metrics endpoint using the Prometheus '
        'client library. Key metrics include request latency (p50, p95, p99), '
        'error rates, and throughput per service.'
    )

    doc.add_heading('Alert Configuration', level=2)

    doc.add_paragraph(
        'Alerts are configured in Prometheus Alertmanager with severity levels: '
        'critical (pages on-call engineer immediately), warning (sends Slack '
        'notification), and info (logged for review). Critical alerts include: '
        'error rate above 5%, p99 latency exceeding 2 seconds, and pod restart '
        'count exceeding 3 in 10 minutes.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
