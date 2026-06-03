"""
Initial Setup: Create a Writer document with a reference guide layout (single-column, no frames).
Task ID: writer_rd_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_072'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('Python Developer Reference Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Version 3.2 - Comprehensive API Reference')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    doc.add_paragraph('')  # spacer

    # Navigation labels as regular paragraphs (single-column, no frames)
    nav_header = doc.add_paragraph()
    run = nav_header.add_run('Table of Contents')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    chapters = [
        ('Chapter 1', 'Getting Started'),
        ('Chapter 2', 'Data Structures'),
        ('Chapter 3', 'Functions & Decorators'),
        ('Chapter 4', 'Object-Oriented Programming'),
        ('Chapter 5', 'Error Handling'),
        ('Chapter 6', 'File I/O Operations'),
        ('Chapter 7', 'Concurrency & Parallelism'),
        ('Chapter 8', 'Testing & Debugging'),
    ]

    for ch_num, ch_title in chapters:
        p = doc.add_paragraph()
        run = p.add_run(f'{ch_num}: {ch_title}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

    doc.add_paragraph('')  # spacer

    # Chapter 1: Getting Started
    h1 = doc.add_heading('Chapter 1: Getting Started', level=1)

    doc.add_paragraph(
        'Python is a versatile, high-level programming language known for its '
        'readability and extensive standard library. This chapter covers the '
        'fundamental setup procedures, environment configuration, and best '
        'practices for starting a new Python project.'
    )

    doc.add_paragraph(
        'Before writing any code, ensure your development environment is '
        'properly configured. Install Python 3.10 or later from the official '
        'website (python.org). Use virtual environments to isolate project '
        'dependencies and avoid version conflicts across different projects.'
    )

    doc.add_paragraph(
        'The recommended project structure includes a src/ directory for '
        'source code, tests/ for unit tests, docs/ for documentation, and '
        'a requirements.txt or pyproject.toml file for dependency management. '
        'Initialize version control with Git from the project root.'
    )

    # Chapter 2: Data Structures
    doc.add_heading('Chapter 2: Data Structures', level=1)

    doc.add_paragraph(
        'Python provides several built-in data structures that are essential '
        'for everyday programming. Lists, tuples, dictionaries, and sets each '
        'serve different purposes and offer distinct performance characteristics. '
        'Understanding when to use each type is crucial for writing efficient code.'
    )

    doc.add_paragraph(
        'Lists are ordered, mutable sequences that support heterogeneous elements. '
        'Common operations include append(), extend(), insert(), and pop(). '
        'List comprehensions provide a concise way to create filtered or '
        'transformed lists: filtered = [x for x in data if x > threshold].'
    )

    doc.add_paragraph(
        'Dictionaries map hashable keys to arbitrary values with O(1) average '
        'lookup time. Since Python 3.7, dictionaries maintain insertion order. '
        'The collections module extends built-in types with defaultdict, '
        'OrderedDict, Counter, and ChainMap for specialized use cases.'
    )

    # Chapter 3: Functions & Decorators
    doc.add_heading('Chapter 3: Functions & Decorators', level=1)

    doc.add_paragraph(
        'Functions are first-class objects in Python, meaning they can be '
        'assigned to variables, passed as arguments, and returned from other '
        'functions. This enables powerful functional programming patterns '
        'including closures, higher-order functions, and decorator chains.'
    )

    doc.add_paragraph(
        'Decorators wrap functions to extend their behavior without modifying '
        'the original source code. Common use cases include logging, caching '
        '(@functools.lru_cache), access control, and performance measurement. '
        'Class-based decorators implement __call__ for stateful decoration.'
    )

    # Chapter 4: Object-Oriented Programming
    doc.add_heading('Chapter 4: Object-Oriented Programming', level=1)

    doc.add_paragraph(
        'Python supports multiple inheritance, abstract base classes, and a '
        'flexible attribute lookup system through the Method Resolution Order '
        '(MRO). The dataclasses module (Python 3.7+) automates boilerplate '
        'for classes that primarily store data, generating __init__, __repr__, '
        'and comparison methods automatically.'
    )

    doc.add_paragraph(
        'Properties (@property) provide controlled access to instance attributes. '
        'Descriptors offer a more general mechanism for attribute access control '
        'by implementing __get__, __set__, and __delete__ methods. Metaclasses '
        'control class creation itself and are used in frameworks like Django ORM.'
    )

    # Chapter 5: Error Handling
    doc.add_heading('Chapter 5: Error Handling', level=1)

    doc.add_paragraph(
        'Robust error handling uses try/except/else/finally blocks to manage '
        'exceptional conditions. Custom exception hierarchies allow fine-grained '
        'error classification. The contextlib module provides context managers '
        'for resource cleanup via @contextmanager decorator and ExitStack.'
    )

    # Chapter 6: File I/O Operations
    doc.add_heading('Chapter 6: File I/O Operations', level=1)

    doc.add_paragraph(
        'Python offers multiple approaches to file handling: built-in open() '
        'with context managers, pathlib.Path for object-oriented path '
        'manipulation, and specialized modules for CSV, JSON, XML, and binary '
        'formats. The io module supports StringIO and BytesIO for in-memory '
        'file-like objects useful in testing and data transformation.'
    )

    # Chapter 7: Concurrency & Parallelism
    doc.add_heading('Chapter 7: Concurrency & Parallelism', level=1)

    doc.add_paragraph(
        'The asyncio framework enables cooperative multitasking with async/await '
        'syntax, ideal for I/O-bound workloads. The threading module provides '
        'OS-level threads constrained by the GIL, while multiprocessing bypasses '
        'the GIL for CPU-bound parallelism. concurrent.futures offers a unified '
        'executor interface for both thread and process pools.'
    )

    # Chapter 8: Testing & Debugging
    doc.add_heading('Chapter 8: Testing & Debugging', level=1)

    doc.add_paragraph(
        'The pytest framework is the de facto standard for Python testing, '
        'supporting fixtures, parameterized tests, and plugin extensions. '
        'Coverage.py measures code coverage and integrates with CI/CD pipelines. '
        'The pdb debugger supports breakpoints, stepping, and variable inspection '
        'directly in the terminal.'
    )

    doc.add_paragraph(
        'Type checking with mypy catches type errors statically before runtime. '
        'Property-based testing with Hypothesis generates edge-case inputs '
        'automatically. Integration tests should use Docker containers or '
        'testcontainers-python to isolate external service dependencies.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
