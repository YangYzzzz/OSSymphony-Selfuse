"""
Initial Setup: Corporate Manual with default TOC title
Task ID: writer_mt_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_058'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Document Title ---
    title = doc.add_heading('Riverside Technologies Corporate Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Version 3.2 | Effective Date: January 15, 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # --- Table of Contents (simulated) ---
    # TOC title - default style: "Table of Contents", normal formatting
    toc_title = doc.add_paragraph()
    toc_title_run = toc_title.add_run('Table of Contents')
    toc_title_run.font.size = Pt(14)
    # Default: NOT bold, NOT centered, NOT all caps - this is what the task should change

    # TOC entries (simulated as plain text with tab leaders)
    toc_entries = [
        ('1. Introduction', '3'),
        ('2. Company Overview', '5'),
        ('3. Employment Policies', '8'),
        ('   3.1 Hiring Process', '8'),
        ('   3.2 Onboarding Procedures', '10'),
        ('   3.3 Performance Reviews', '12'),
        ('4. Code of Conduct', '15'),
        ('5. Benefits and Compensation', '18'),
        ('   5.1 Health Insurance', '18'),
        ('   5.2 Retirement Plans', '20'),
        ('   5.3 Paid Time Off', '22'),
        ('6. Safety and Security', '25'),
        ('7. IT Policies', '28'),
        ('8. Appendices', '32'),
    ]
    for entry_text, page_num in toc_entries:
        p = doc.add_paragraph()
        r = p.add_run(f'{entry_text} {"." * (50 - len(entry_text))} {page_num}')
        r.font.size = Pt(11)

    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Welcome to Riverside Technologies. This corporate manual serves as a comprehensive '
        'guide to our company policies, procedures, and expectations. All employees are expected '
        'to familiarize themselves with the contents of this document and adhere to the guidelines '
        'outlined herein.'
    )
    doc.add_paragraph(
        'Riverside Technologies was founded in 2008 with a mission to deliver innovative software '
        'solutions that empower businesses worldwide. Our commitment to excellence drives everything '
        'we do, from product development to customer support.'
    )
    doc.add_paragraph(
        'This manual is reviewed annually by the Human Resources department and the Executive '
        'Leadership Team. The most recent revision was approved on December 12, 2024, by Chief '
        'Operating Officer Elena Vasquez.'
    )

    doc.add_page_break()

    # --- Chapter 2: Company Overview ---
    doc.add_heading('2. Company Overview', level=1)
    doc.add_paragraph(
        'Riverside Technologies operates in 14 countries with over 3,200 employees. Our headquarters '
        'is located in Austin, Texas, with regional offices in London, Singapore, and Sao Paulo.'
    )

    doc.add_heading('Our Core Values', level=2)
    doc.add_paragraph('Innovation - We embrace creative thinking and continuous improvement.', style='List Bullet')
    doc.add_paragraph('Integrity - We conduct business with honesty and transparency.', style='List Bullet')
    doc.add_paragraph('Collaboration - We believe in the power of teamwork.', style='List Bullet')
    doc.add_paragraph('Customer Focus - We prioritize the needs of our clients.', style='List Bullet')
    doc.add_paragraph('Sustainability - We are committed to environmental responsibility.', style='List Bullet')

    doc.add_page_break()

    # --- Chapter 3: Employment Policies ---
    doc.add_heading('3. Employment Policies', level=1)

    doc.add_heading('3.1 Hiring Process', level=2)
    doc.add_paragraph(
        'All open positions must be approved by the department head and the VP of Human Resources '
        'before posting. Candidates undergo a structured interview process consisting of a phone '
        'screening, technical assessment, and panel interview.'
    )

    doc.add_heading('3.2 Onboarding Procedures', level=2)
    doc.add_paragraph(
        'New hires participate in a two-week onboarding program that includes orientation sessions, '
        'IT setup, compliance training, and introductions to key team members. Each new employee is '
        'assigned a mentor from their department for the first 90 days.'
    )

    doc.add_heading('3.3 Performance Reviews', level=2)
    doc.add_paragraph(
        'Performance evaluations are conducted semi-annually in June and December. Managers use '
        'the standardized evaluation framework, which assesses competencies across five dimensions: '
        'technical skills, communication, leadership, initiative, and teamwork.'
    )

    doc.add_page_break()

    # --- Chapter 4: Code of Conduct ---
    doc.add_heading('4. Code of Conduct', level=1)
    doc.add_paragraph(
        'All employees are expected to maintain the highest standards of professional conduct. '
        'This includes respecting colleagues, protecting company assets, and complying with all '
        'applicable laws and regulations.'
    )
    doc.add_paragraph(
        'Violations of the code of conduct may result in disciplinary action, up to and including '
        'termination of employment. Employees are encouraged to report concerns through the '
        'confidential ethics hotline at 1-800-555-ETHICS.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
