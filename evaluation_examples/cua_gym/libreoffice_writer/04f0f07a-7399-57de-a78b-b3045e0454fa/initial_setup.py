"""
Initial Setup: Create a 3-page business letter document with no headers or footers.
Task ID: writer_pd_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page to standard US Letter with default margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Ensure NO header/footer
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    # Clear any default header/footer content
    for p in section.header.paragraphs:
        p.clear()
    for p in section.footer.paragraphs:
        p.clear()

    # --- PAGE 1: Company Letter ---
    heading = doc.add_heading('Meridian Dynamics', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    date_para = doc.add_paragraph('April 1, 2026')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph('')  # blank line

    doc.add_paragraph('Ms. Elena Rodriguez')
    doc.add_paragraph('Director of Operations')
    doc.add_paragraph('Pacific Northwest Regional Office')
    doc.add_paragraph('4200 Cascade Boulevard, Suite 310')
    doc.add_paragraph('Portland, OR 97201')

    doc.add_paragraph('')

    greeting = doc.add_paragraph('Dear Ms. Rodriguez,')

    doc.add_paragraph(
        'I am writing to confirm the details of our upcoming quarterly strategy review '
        'meeting scheduled for April 15, 2026. As discussed during our last conference call, '
        'this meeting will focus on the operational performance of Q1 2026 and outline our '
        'strategic priorities for the remainder of the fiscal year.'
    )

    doc.add_paragraph(
        'The Pacific Northwest region has shown remarkable growth over the past quarter, '
        'with a 12.3% increase in client acquisitions and a 7.8% improvement in customer '
        'retention rates. These figures exceed our initial projections by a significant margin, '
        'and the executive team is eager to understand the factors driving this success.'
    )

    doc.add_paragraph(
        'Please prepare a detailed presentation covering the following key areas for the '
        'upcoming review session:'
    )

    # Bullet list items
    items = [
        'Revenue breakdown by product line and service category for Q1 2026',
        'Client acquisition funnel metrics and conversion rate analysis',
        'Staffing utilization rates and resource allocation efficiency',
        'Customer satisfaction survey results and Net Promoter Score trends',
        'Competitive landscape assessment for the Pacific Northwest market',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'In addition to the presentation materials, please compile a one-page executive '
        'summary that highlights the top three achievements and the top three challenges '
        'your team encountered during Q1. This summary will be distributed to the board '
        'of directors as part of our quarterly reporting package.'
    )

    # --- PAGE 2 (force page break) ---
    doc.add_page_break()

    doc.add_heading('Meeting Logistics', level=2)

    doc.add_paragraph(
        'The quarterly review will be held at our corporate headquarters in San Francisco, '
        'located at 1 Market Street, 28th Floor. The meeting is scheduled to begin at '
        '9:00 AM Pacific Time and is expected to conclude by 4:00 PM. A working lunch will '
        'be provided in the executive conference room.'
    )

    doc.add_paragraph(
        'For those joining remotely, a video conference link will be distributed by '
        'our IT department no later than April 12, 2026. Please ensure your team members '
        'test their connections in advance to avoid any technical delays during the session.'
    )

    doc.add_paragraph('The agenda for the day is as follows:')

    agenda_items = [
        '9:00 AM - 9:30 AM: Welcome and opening remarks by CEO, David Nakamura',
        '9:30 AM - 10:45 AM: Financial performance overview (CFO, Margaret Liu)',
        '10:45 AM - 11:00 AM: Coffee break',
        '11:00 AM - 12:15 PM: Regional performance presentations',
        '12:15 PM - 1:15 PM: Working lunch and informal discussion',
        '1:15 PM - 2:30 PM: Strategic initiative updates',
        '2:30 PM - 2:45 PM: Afternoon break',
        '2:45 PM - 3:45 PM: Open forum and Q&A session',
        '3:45 PM - 4:00 PM: Closing remarks and action item review',
    ]
    for item in agenda_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_paragraph(
        'Please note that the regional performance presentations are limited to 15 minutes '
        'each, followed by a 5-minute Q&A period. We have six regions presenting, so '
        'adherence to the time constraints is essential to keep the meeting on schedule.'
    )

    doc.add_paragraph(
        'Travel arrangements for out-of-town participants should be coordinated through '
        'our corporate travel desk. Contact Amanda Prescott at extension 4471 or via email '
        'at a.prescott@meridiandynamics.com to arrange flights and hotel accommodations. '
        'The preferred hotel is the Hyatt Regency Embarcadero, where we have negotiated '
        'a corporate rate of $189 per night.'
    )

    # --- PAGE 3 (force page break) ---
    doc.add_page_break()

    doc.add_heading('Action Items and Follow-Up', level=2)

    doc.add_paragraph(
        'Following the quarterly review, each regional director will be expected to submit '
        'an updated operational plan for Q2 2026 within two weeks of the meeting. The plan '
        'should address any gaps identified during the review and incorporate feedback '
        'received from the executive team.'
    )

    doc.add_paragraph(
        'Key deliverables expected from your team include:'
    )

    deliverables = [
        'Revised revenue forecast for Q2-Q4 2026 with monthly granularity',
        'Staffing plan addressing the three open positions in your region',
        'Client retention risk assessment and mitigation strategy',
        'Technology infrastructure upgrade timeline and budget estimate',
        'Partnership development pipeline with probability-weighted projections',
    ]
    for item in deliverables:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Our goal for fiscal year 2026 remains ambitious: achieving $47.5 million in total '
        'revenue across all regions while maintaining our operating margin above 18%. '
        'The strong Q1 performance gives us confidence that these targets are achievable, '
        'but continued discipline in cost management and strategic investment will be critical.'
    )

    doc.add_paragraph(
        'If you have any questions or require additional information prior to the meeting, '
        'please do not hesitate to reach out to my office directly. I look forward to '
        'a productive session and to celebrating the accomplishments of your outstanding team.'
    )

    doc.add_paragraph('')
    doc.add_paragraph('Warm regards,')
    doc.add_paragraph('')
    doc.add_paragraph('Jonathan Whitfield')
    doc.add_paragraph('Senior Vice President, Strategic Operations')
    doc.add_paragraph('Meridian Dynamics Corporation')
    doc.add_paragraph('Phone: (415) 555-0192')
    doc.add_paragraph('Email: j.whitfield@meridiandynamics.com')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
