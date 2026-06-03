"""
Reward Script: Branded header and footer for Company_Letter.docx
Task ID: writer_pd_010
Domain: libreoffice_writer
Scoring:
  Component 1: Header text "Meridian Dynamics" present and left-aligned (0.15)
  Component 2: Header font: bold, 14pt, blue #003366 (0.20)
  Component 3: Header bottom border / horizontal line (0.20)
  Component 4: Footer contains PAGE field code (0.15)
  Component 5: Footer contains NUMPAGES field code (0.15)
  Component 6: Footer contains date "2026-04-01" (0.15)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_010'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def persist_app_state():
    """Best-effort save via Ctrl+S in case doc is still open in LibreOffice."""
    try:
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        import time
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) == 0:
        print("CRITICAL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    header = section.header
    footer = section.footer

    # Component 1: Header text "Meridian Dynamics" present and left-aligned (0.15 points)
    try:
        header_text = ""
        header_alignment = None
        if header.paragraphs:
            header_text = header.paragraphs[0].text.strip()
            header_alignment = header.paragraphs[0].paragraph_format.alignment

        has_correct_text = "Meridian Dynamics" in header_text
        # Left-aligned means alignment is LEFT (0) or None (default=left)
        is_left = header_alignment is None or header_alignment == WD_PARAGRAPH_ALIGNMENT.LEFT

        if has_correct_text and is_left:
            print(f"PASS: Component 1 - Header has 'Meridian Dynamics' left-aligned (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 - Header text='{header_text}', alignment={header_alignment}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Header font formatting: bold, 14pt, blue #003366 (0.20 points)
    try:
        found_correct_format = False
        if header.paragraphs:
            for run in header.paragraphs[0].runs:
                if "Meridian Dynamics" in run.text:
                    is_bold = run.font.bold is True
                    is_14pt = run.font.size is not None and abs(run.font.size.pt - 14.0) < 0.5
                    has_blue = (run.font.color.rgb is not None and
                                str(run.font.color.rgb).upper() == '003366')

                    if is_bold and is_14pt and has_blue:
                        found_correct_format = True
                        print(f"PASS: Component 2 - Header font: bold={is_bold}, size={run.font.size.pt}pt, color={run.font.color.rgb} (0.20 pts)")
                    else:
                        print(f"FAIL: Component 2 - bold={is_bold}, 14pt={is_14pt} (size={run.font.size}), blue={has_blue} (color={run.font.color.rgb})")
                    break

        if found_correct_format:
            total_score += 0.20
        elif not header.paragraphs or not header.paragraphs[0].runs:
            print(f"FAIL: Component 2 - No runs found in header paragraph")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Header bottom border (horizontal line below header) (0.20 points)
    try:
        has_bottom_border = False
        if header.paragraphs:
            para_elem = header.paragraphs[0]._element
            pPr = para_elem.find('w:pPr', NS)
            if pPr is not None:
                pBdr = pPr.find('w:pBdr', NS)
                if pBdr is not None:
                    bottom = pBdr.find('w:bottom', NS)
                    if bottom is not None:
                        val = bottom.get(qn('w:val'))
                        # Any non-none border value indicates a line
                        if val and val != 'none':
                            has_bottom_border = True

        if has_bottom_border:
            print(f"PASS: Component 3 - Header has bottom border (horizontal line) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 - No bottom border found on header paragraph")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Footer contains PAGE field code (0.15 points)
    try:
        has_page_field = False
        if footer.paragraphs:
            for para in footer.paragraphs:
                fields = para._element.findall('.//w:instrText', NS)
                for f in fields:
                    if f.text and 'PAGE' in f.text and 'NUMPAGES' not in f.text:
                        has_page_field = True
                        break

        if has_page_field:
            print(f"PASS: Component 4 - Footer has PAGE field code (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 - No PAGE field code found in footer")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Footer contains NUMPAGES field code (0.15 points)
    try:
        has_numpages_field = False
        if footer.paragraphs:
            for para in footer.paragraphs:
                fields = para._element.findall('.//w:instrText', NS)
                for f in fields:
                    if f.text and 'NUMPAGES' in f.text:
                        has_numpages_field = True
                        break

        if has_numpages_field:
            print(f"PASS: Component 5 - Footer has NUMPAGES field code (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 - No NUMPAGES field code found in footer")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Footer contains date "2026-04-01" (0.15 points)
    try:
        has_date = False
        if footer.paragraphs:
            for para in footer.paragraphs:
                if '2026-04-01' in para.text:
                    has_date = True
                    break

        if has_date:
            print(f"PASS: Component 6 - Footer contains date '2026-04-01' (0.15 pts)")
            total_score += 0.15
        else:
            footer_texts = [p.text for p in footer.paragraphs] if footer.paragraphs else []
            print(f"FAIL: Component 6 - Date '2026-04-01' not found in footer. Footer texts: {footer_texts}")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
