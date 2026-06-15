"""
Initial Setup: Insert soft hyphen in 'internationalization' in paragraph 2
Task ID: writer_txtfmt_066
Domain: libreoffice_writer

Creates localization_guide.docx on the Desktop with 'internationalization'
having NO soft hyphen (pre-task state).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_066'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/localization_guide.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # --- Paragraph 1: Title / Introduction heading ---
    heading = doc.add_heading('Software Localization Guide', level=1)

    # --- Paragraph 1 (body): Introduction ---
    intro = doc.add_paragraph(
        'This guide provides an overview of localization best practices for modern software '
        'development teams. It covers key concepts, workflows, and tools used to bring products '
        'to global markets efficiently and consistently.'
    )

    # --- Paragraph 2: Contains 'internationalization' WITHOUT soft hyphen ---
    # NOTE: 'internationalization' here has NO soft hyphen — that is the pre-task state.
    para2 = doc.add_paragraph(
        'The process of internationalization involves designing software so it can be adapted '
        'to various languages and regions without engineering changes. '
        'Teams that embrace internationalization early in the development lifecycle tend to '
        'reduce the cost and complexity of future localization efforts significantly.'
    )

    # --- Paragraph 3: Technical details ---
    doc.add_paragraph(
        'Key components of a robust localization framework include Unicode support, locale-aware '
        'date and time formatting, currency handling, and right-to-left (RTL) text rendering. '
        'Each component must be validated against a comprehensive set of locale-specific test cases.'
    )

    # --- Paragraph 4: Process overview ---
    doc.add_paragraph(
        'The localization workflow typically begins with string extraction from the source code, '
        'followed by translation by professional linguists, and then integration testing in the '
        'target locale. Automated regression suites help catch regressions introduced during '
        'the translation and integration phases.'
    )

    # --- Paragraph 5: Conclusion ---
    doc.add_paragraph(
        'Investing in proper internationalization infrastructure pays dividends as the product '
        'scales to new markets. Organizations that treat localization as a first-class engineering '
        'concern consistently achieve faster time-to-market for regional releases.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
