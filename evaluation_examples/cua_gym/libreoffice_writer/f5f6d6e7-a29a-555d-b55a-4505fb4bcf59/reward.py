"""
Reward Script: Insert soft hyphen in 'internationalization' in paragraph 2
Task ID: writer_txtfmt_066
Domain: libreoffice_writer
Scoring:
  Component 1 (0.7): Soft hyphen U+00AD present somewhere in paragraph 2
  Component 2 (0.3): Soft hyphen placed specifically between 'inter' and 'nationalization'
                     in the first occurrence of 'internationalization' in paragraph 2

Both components FAIL on initial_env (no soft hyphen present) and PASS on golden_env.
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
FILE_PATH = f'{WORKDIR}/localization_guide.docx'

SOFT_HYPHEN = '\u00ad'  # Unicode U+00AD soft hyphen

# The target word with soft hyphen at the correct position
EXPECTED_WORD = 'inter\u00adnationalization'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: paragraph 2 must exist (index 2, 0-based)
    if len(doc.paragraphs) < 3:
        print(f"CRITICAL: Document has fewer than 3 paragraphs (found {len(doc.paragraphs)})")
        print("REWARD: 0.0")
        return 0.0

    para2 = doc.paragraphs[2]
    para2_text = para2.text  # all runs concatenated

    # Component 1: A soft hyphen (U+00AD) is present somewhere in paragraph 2 (0.7 points)
    # This FAILS on initial_env (no soft hyphen anywhere in para 2) and PASSES on golden_env.
    try:
        if SOFT_HYPHEN in para2_text:
            print(f"PASS: Component 1 — soft hyphen U+00AD present in paragraph 2 (0.7 pts)")
            total_score += 0.7
        else:
            print(f"FAIL: Component 1 — no soft hyphen (U+00AD) found in paragraph 2")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: The soft hyphen is at the correct position — between 'inter' and 'nationalization'
    # i.e. the string 'inter<SH>nationalization' appears in paragraph 2 (0.3 points)
    # This FAILS on initial_env and PASSES only on golden_env with the correct insertion point.
    try:
        if EXPECTED_WORD in para2_text:
            total_score += 0.3
            idx = para2_text.find(EXPECTED_WORD)
            context = para2_text[max(0, idx - 5):idx + len(EXPECTED_WORD) + 5]
            print(f"PASS: Component 2 — 'inter<SH>nationalization' found at index {idx}: {repr(context)} (0.3 pts)")
        elif SOFT_HYPHEN in para2_text:
            pos = para2_text.find(SOFT_HYPHEN)
            context = para2_text[max(0, pos - 5):pos + 15]
            print(f"FAIL: Component 2 — soft hyphen at wrong position in para 2, context: {repr(context)}")
            print(f"  Expected: 'inter<SH>nationalization', found hyphen elsewhere")
        else:
            print(f"FAIL: Component 2 — no soft hyphen found in paragraph 2")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
