"""
Reward Script: Create a 10-row by 3-column guest list table in LibreOffice Writer
Task ID: writer_tbl_053
Domain: libreoffice_writer

Scoring Rubric (total = 1.0):
  Component 1: Table structure (10 rows x 3 cols, correct header, 9 guest names)  — 0.40
  Component 2: Header row repeats across pages (tblHeader = 1)                    — 0.20
  Component 3: Alternating row shading (odd data rows D3D3D3, even FFFFFF)         — 0.20
  Component 4: Table centered on page (jc = center)                               — 0.20
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_053'
FILE_PATH = f'{WORKDIR}/Desktop/guest_list.docx'

HEADER_COLS = ['Guest Name', 'RSVP', 'Dietary Needs']
EXPECTED_ROWS = 10
EXPECTED_COLS = 3


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must exist and be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: there must be at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document (precondition failed)")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # -------------------------------------------------------------------------
    # Component 1: Table structure — 10 rows x 3 cols, correct header, 9 guest names (0.40 pts)
    # -------------------------------------------------------------------------
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        row0_cells = [table.cell(0, j).text.strip() for j in range(num_cols)]

        structure_ok = (num_rows == EXPECTED_ROWS and num_cols == EXPECTED_COLS)
        header_ok = (row0_cells == HEADER_COLS)

        # Count non-empty guest names in column 1, rows 1-9
        guest_names = []
        for r in range(1, num_rows):
            name = table.cell(r, 0).text.strip()
            if name:
                guest_names.append(name)
        names_ok = len(guest_names) == 9

        if structure_ok and header_ok and names_ok:
            print(f"PASS: Component 1 — Table is {num_rows}x{num_cols}, header={row0_cells}, "
                  f"9 guest names found (0.40 pts)")
            total_score += 0.40
        else:
            details = []
            if not structure_ok:
                details.append(f"expected 10x3 got {num_rows}x{num_cols}")
            if not header_ok:
                details.append(f"header mismatch: expected {HEADER_COLS} got {row0_cells}")
            if not names_ok:
                details.append(f"expected 9 guest names, found {len(guest_names)}: {guest_names}")
            print(f"FAIL: Component 1 — {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Header row repeats across pages (tblHeader element present on row 0) (0.20 pts)
    # -------------------------------------------------------------------------
    try:
        header_row = table.rows[0]
        tr_pr = header_row._tr.find(qn('w:trPr'))
        tbl_header_found = False
        if tr_pr is not None:
            tbl_header = tr_pr.find(qn('w:tblHeader'))
            if tbl_header is not None:
                val = tbl_header.get(qn('w:val'))
                # val='1' or val='true' or None all indicate the header repeats
                if val in ('1', 'true', None):
                    tbl_header_found = True

        if tbl_header_found:
            print("PASS: Component 2 — Header row repeat (tblHeader=1) is set (0.20 pts)")
            total_score += 0.20
        else:
            print("FAIL: Component 2 — Header row does not have tblHeader repeat property set")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Alternating row shading — odd data rows light gray (D3D3D3),
    # even data rows white/none (FFFFFF or no shading) (0.20 pts)
    # -------------------------------------------------------------------------
    try:
        # Data rows are rows 1-9 (0-indexed). Odd data rows: rows 1,3,5,7,9 → should be gray.
        # Even data rows: rows 2,4,6,8 → should be white/default.
        GRAY_FILL = 'D3D3D3'
        WHITE_FILLS = ('FFFFFF', None)  # None means no shading element

        gray_rows_ok = True
        white_rows_ok = True
        details_gray = []
        details_white = []

        for r_idx in range(1, EXPECTED_ROWS):
            # Determine expected: odd row indices (1,3,5,7,9) should be gray
            expected_gray = (r_idx % 2 == 1)  # rows 1,3,5,7,9
            row = table.rows[r_idx]
            row_fills = []
            for c_idx in range(EXPECTED_COLS):
                cell = table.cell(r_idx, c_idx)
                tc_pr = cell._tc.find(qn('w:tcPr'))
                fill = None
                if tc_pr is not None:
                    shd = tc_pr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                row_fills.append(fill)

            # All cells in the row should have the same fill
            unique_fills = set(row_fills)
            actual_fill = row_fills[0] if row_fills else None

            if expected_gray:
                if actual_fill and actual_fill.upper() == GRAY_FILL:
                    pass  # OK
                else:
                    gray_rows_ok = False
                    details_gray.append(f"row {r_idx} expected D3D3D3 got {actual_fill}")
            else:
                if actual_fill is None or actual_fill.upper() in ('FFFFFF', 'AUTO'):
                    pass  # OK
                else:
                    white_rows_ok = False
                    details_white.append(f"row {r_idx} expected FFFFFF/None got {actual_fill}")

        if gray_rows_ok and white_rows_ok:
            print("PASS: Component 3 — Alternating row shading correct (odd=D3D3D3, even=FFFFFF) (0.20 pts)")
            total_score += 0.20
        else:
            issues = details_gray + details_white
            print(f"FAIL: Component 3 — Alternating shading issues: {issues}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Table centered on page (jc = center in tblPr) (0.20 pts)
    # -------------------------------------------------------------------------
    try:
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        centered = False
        if tbl_pr is not None:
            jc = tbl_pr.find(qn('w:jc'))
            if jc is not None:
                val = jc.get(qn('w:val'))
                if val and val.lower() == 'center':
                    centered = True

        if centered:
            print("PASS: Component 4 — Table is centered on the page (jc=center) (0.20 pts)")
            total_score += 0.20
        else:
            tbl_pr_xml = tbl_pr is not None
            jc_val = None
            if tbl_pr is not None:
                jc = tbl_pr.find(qn('w:jc'))
                if jc is not None:
                    jc_val = jc.get(qn('w:val'))
            print(f"FAIL: Component 4 — Table not centered; tblPr found={tbl_pr_xml}, jc val={jc_val}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
