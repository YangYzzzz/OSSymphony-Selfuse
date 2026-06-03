"""
Initial Setup: Employee Skills Inventory with Duplicate Entries
Task ID: osworld_writer_duplicate_line_removal_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_duplicate_line_removal_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title heading
    title = doc.add_heading('Employee Skills Inventory', level=1)

    # Introduction paragraph
    intro = doc.add_paragraph(
        'This document catalogs the technical and soft skills held by members of the '
        'Horizon Technologies workforce. Skills are organized as reported by department '
        'heads during the Q1 2025 skills assessment.'
    )

    # Section heading
    doc.add_heading('Identified Employee Skills', level=2)

    # 20 skill lines — 13 unique skills with 7 duplicates interspersed
    skills = [
        'Python Programming',           # 1 — unique first occurrence
        'Data Analysis',                # 2 — unique first occurrence
        'Project Management',           # 3 — unique first occurrence
        'SQL Database Management',      # 4 — unique first occurrence
        'Machine Learning',             # 5 — unique first occurrence
        'Communication Skills',         # 6 — unique first occurrence
        'Team Leadership',              # 7 — unique first occurrence
        'Python Programming',           # 8 — DUPLICATE of line 1
        'Cloud Computing (AWS)',        # 9 — unique first occurrence
        'JavaScript Development',       # 10 — unique first occurrence
        'Team Leadership',              # 11 — DUPLICATE of line 7
        'Financial Reporting',          # 12 — unique first occurrence
        'Problem Solving',              # 13 — unique first occurrence
        'Data Analysis',                # 14 — DUPLICATE of line 2
        'Agile Methodology',            # 15 — unique first occurrence
        'Communication Skills',         # 16 — DUPLICATE of line 6
        'Machine Learning',             # 17 — DUPLICATE of line 5
        'Customer Relations',           # 18 — unique first occurrence
        'SQL Database Management',      # 19 — DUPLICATE of line 4
        'Project Management',           # 20 — DUPLICATE of line 3
    ]

    for skill in skills:
        doc.add_paragraph(skill)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
