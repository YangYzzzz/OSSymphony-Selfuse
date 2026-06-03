"""
Reward Script: Create signature blocks at the bottom of NDA
Task ID: writer_hr_032
Domain: libreoffice_writer
Scoring:
  Component 1: Signature table exists with 2 columns (0.25)
  Component 2: Column headers 'Employee' and 'Company Representative' (0.25)
  Component 3: Four labeled line fields per column (Signature, Printed Name, Title, Date) (0.35)
  Component 4: Witness/execution clause paragraph added (0.15)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_032'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            import time
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that signature blocks were added to the NDA document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Signature table exists with 2 columns (0.25 points)
    # The task requires two signature sections side by side, implemented as a table
    try:
        sig_table = None
        for table in doc.tables:
            if len(table.columns) >= 2:
                # Check if this table contains signature-related content
                all_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text.lower() + " "
                if "signature" in all_text or "employee" in all_text:
                    sig_table = table
                    break

        if sig_table is not None and len(sig_table.columns) >= 2:
            print(f"PASS: Component 1 -- Signature table found with {len(sig_table.columns)} columns, {len(sig_table.rows)} rows (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 -- No signature table with 2+ columns found. Tables in doc: {len(doc.tables)}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Column headers 'Employee' and 'Company Representative' (0.25 points)
    # The task specifies two sections: one for Employee, one for Company Representative
    try:
        if sig_table is not None:
            # Collect all cell texts to find the headers
            all_cell_texts = []
            for row in sig_table.rows:
                for cell in row.cells:
                    all_cell_texts.append(cell.text.strip().lower())

            has_employee = any("employee" in t for t in all_cell_texts)
            has_company_rep = any("company representative" in t for t in all_cell_texts)

            if has_employee and has_company_rep:
                print(f"PASS: Component 2 -- Found 'Employee' and 'Company Representative' headers (0.25 pts)")
                total_score += 0.25
            elif has_employee or has_company_rep:
                found = "Employee" if has_employee else "Company Representative"
                missing = "Company Representative" if has_employee else "Employee"
                print(f"PARTIAL: Component 2 -- Found '{found}' but missing '{missing}' (0.125 pts)")
                total_score += 0.125
            else:
                print(f"FAIL: Component 2 -- Neither 'Employee' nor 'Company Representative' found in table cells")
        else:
            print(f"FAIL: Component 2 -- No signature table to check headers")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Four labeled line fields per column (0.35 points)
    # Each column should have: Signature, Printed Name, Title, Date
    try:
        if sig_table is not None:
            required_labels = ["signature", "printed name", "title", "date"]

            # Gather text from each column
            col_count = len(sig_table.columns)
            col_texts = [[] for _ in range(col_count)]
            for row in sig_table.rows:
                for ci, cell in enumerate(row.cells):
                    if ci < col_count:
                        col_texts[ci].append(cell.text.strip().lower())

            # Check first two columns for the required labels
            col0_labels_found = 0
            col1_labels_found = 0
            for label in required_labels:
                # Check column 0
                if any(label in t for t in col_texts[0]):
                    col0_labels_found += 1
                # Check column 1 (if exists)
                if col_count >= 2 and any(label in t for t in col_texts[1]):
                    col1_labels_found += 1

            total_labels = col0_labels_found + col1_labels_found
            max_labels = 8  # 4 per column x 2 columns
            label_ratio = total_labels / max_labels

            if total_labels == max_labels:
                print(f"PASS: Component 3 -- All 4 labels found in both columns ({total_labels}/{max_labels}) (0.35 pts)")
                total_score += 0.35
            elif total_labels > 0:
                label_score = 0.35 * label_ratio
                print(f"PARTIAL: Component 3 -- {total_labels}/{max_labels} labels found (col0={col0_labels_found}/4, col1={col1_labels_found}/4) ({label_score:.3f} pts)")
                total_score += label_score
            else:
                print(f"FAIL: Component 3 -- No signature labels found in table")
        else:
            print(f"FAIL: Component 3 -- No signature table to check labels")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Witness/execution clause paragraph (0.15 points)
    # Golden has "IN WITNESS WHEREOF" paragraph before the signature table
    # This is a new paragraph not present in the initial document
    try:
        witness_paras = [p for p in doc.paragraphs
                         if "in witness whereof" in p.text.strip().lower()
                         or "witness" in p.text.strip().lower()]

        if len(witness_paras) > 0:
            print(f"PASS: Component 4 -- Witness/execution clause paragraph found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 -- No witness/execution clause paragraph found")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
