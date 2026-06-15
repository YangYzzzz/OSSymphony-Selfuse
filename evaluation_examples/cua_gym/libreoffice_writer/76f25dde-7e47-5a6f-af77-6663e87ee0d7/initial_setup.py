"""
Initial Setup: Create NDA document without signature blocks
Task ID: writer_hr_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(12)
    run = preamble.add_run(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of March 15, 2025, '
        'by and between Meridian Technologies, Inc., a Delaware corporation with its principal '
        'offices located at 4500 Innovation Drive, Suite 300, Austin, TX 78759 ("Company"), '
        'and the undersigned employee ("Employee").'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    run = intro.add_run(
        'WHEREAS, the Employee has been offered employment or is currently employed by the Company; and '
        'WHEREAS, the Company possesses certain confidential and proprietary information critical to its '
        'business operations; and WHEREAS, both parties agree that the protection of such information is '
        'essential to the Company\'s competitive position and ongoing success;'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    now_therefore = doc.add_paragraph()
    now_therefore.paragraph_format.space_after = Pt(12)
    run = now_therefore.add_run(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, '
        'and for other good and valuable consideration, the receipt and sufficiency of which are hereby '
        'acknowledged, the parties agree as follows:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 1: Definitions ---
    h1 = doc.add_heading('1. DEFINITIONS', level=1)
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(8)
    run = p1.add_run(
        '"Confidential Information" shall mean any and all non-public information, whether in written, '
        'oral, electronic, or visual form, that is disclosed by the Company to the Employee, including '
        'but not limited to: (a) trade secrets, inventions, patents, copyrights, and other intellectual '
        'property; (b) business plans, strategies, forecasts, and financial data; (c) customer and '
        'supplier lists, pricing information, and market analyses; (d) software source code, algorithms, '
        'system architectures, and technical specifications; (e) personnel information, compensation '
        'data, and organizational plans; and (f) any other information designated as confidential by '
        'the Company or that a reasonable person would understand to be confidential given the nature '
        'of the information and circumstances of disclosure.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 2: Obligations of the Employee ---
    doc.add_heading('2. OBLIGATIONS OF THE EMPLOYEE', level=1)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run(
        'The Employee agrees to: (a) hold all Confidential Information in strict confidence and not '
        'disclose it to any third party without the prior written consent of the Company; (b) use '
        'Confidential Information solely for the purpose of performing duties as an employee of the '
        'Company; (c) take all reasonable precautions to prevent unauthorized disclosure or use of '
        'Confidential Information, including but not limited to implementing appropriate physical and '
        'electronic security measures; (d) promptly notify the Company of any unauthorized disclosure '
        'or use of Confidential Information of which the Employee becomes aware; and (e) not reverse '
        'engineer, decompile, or disassemble any products or materials embodying Confidential Information.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 3: Exclusions ---
    doc.add_heading('3. EXCLUSIONS FROM CONFIDENTIAL INFORMATION', level=1)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    run = p3.add_run(
        'The obligations set forth in this Agreement shall not apply to information that: '
        '(a) is or becomes publicly available through no fault of the Employee; (b) was known to '
        'the Employee prior to its disclosure by the Company, as demonstrated by written records; '
        '(c) is independently developed by the Employee without reference to the Confidential '
        'Information; (d) is rightfully received by the Employee from a third party without '
        'restriction on disclosure; or (e) is required to be disclosed by law, regulation, or '
        'court order, provided that the Employee gives the Company prompt written notice of such '
        'requirement and cooperates with the Company in seeking a protective order.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 4: Term and Termination ---
    doc.add_heading('4. TERM AND TERMINATION', level=1)
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(8)
    run = p4.add_run(
        'This Agreement shall be effective as of the date first written above and shall continue '
        'in full force and effect during the term of the Employee\'s employment with the Company '
        'and for a period of three (3) years following the termination of employment for any reason. '
        'Upon termination of employment, the Employee shall promptly return to the Company all '
        'documents, files, and other materials containing or relating to Confidential Information, '
        'including all copies thereof, whether in physical or electronic form.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 5: Remedies ---
    doc.add_heading('5. REMEDIES', level=1)
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(8)
    run = p5.add_run(
        'The Employee acknowledges that any breach of this Agreement may cause irreparable harm '
        'to the Company for which monetary damages may be inadequate. Accordingly, in addition to '
        'any other remedies available at law or in equity, the Company shall be entitled to seek '
        'injunctive relief and specific performance to enforce the terms of this Agreement without '
        'the necessity of proving actual damages or posting a bond.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 6: Miscellaneous ---
    doc.add_heading('6. MISCELLANEOUS', level=1)
    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(8)
    run = p6.add_run(
        '(a) Governing Law. This Agreement shall be governed by and construed in accordance with '
        'the laws of the State of Delaware, without regard to its conflict of law provisions. '
        '(b) Entire Agreement. This Agreement constitutes the entire agreement between the parties '
        'with respect to the subject matter hereof and supersedes all prior and contemporaneous '
        'understandings, agreements, representations, and warranties. '
        '(c) Amendments. No amendment or modification of this Agreement shall be effective unless '
        'made in writing and signed by both parties. '
        '(d) Severability. If any provision of this Agreement is held to be invalid or unenforceable, '
        'the remaining provisions shall continue in full force and effect. '
        '(e) Assignment. The Employee may not assign or transfer this Agreement without the prior '
        'written consent of the Company. The Company may assign this Agreement to any successor '
        'or affiliate without the Employee\'s consent.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Document ends abruptly here — NO signature blocks

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
