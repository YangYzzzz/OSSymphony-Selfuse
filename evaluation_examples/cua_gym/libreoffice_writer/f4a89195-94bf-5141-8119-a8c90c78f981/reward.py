"""
Reward Script: Add Average row to grade table in class_grades.docx
Task ID: writer_tbl_077
Domain: libreoffice_writer
Scoring:
  Component 1: Table has 7 rows (average row added)          — 0.2 pts
  Component 2: Row 7, Col A = 'Average' label                — 0.2 pts
  Component 3: Row 7, Col B = '86' (mean of 85,92,78,95,80) — 0.2 pts
  Component 4: Row 7, Col C = '82' (mean of 78,88,82,90,72) — 0.2 pts
  Component 5: Row 7, Col D = '86.6' (mean of 90,85,88,92,78)— 0.2 pts
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_077'

def verify_task(file_path):
    """
    Verify that a 7th row was added to the grade table with:
    - A7 labeled 'Average'
    - B7 showing '86' (mean of Homework scores)
    - C7 showing '82' (mean of Midterm scores)
    - D7 showing '86.6' (mean of Final scores)

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — gate: if it can't be loaded, return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gate: Must have exactly 1 table
    if len(doc.tables) < 1:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table has 7 rows (the Average row was added) (0.2 points)
    try:
        num_rows = len(table.rows)
        if num_rows >= 7:
            print(f"PASS: Component 1 — Table has {num_rows} rows (>= 7, Average row added) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 1 — Table has {num_rows} rows, expected at least 7 (Average row missing)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Only proceed with row 6 checks if row exists
    if len(table.rows) < 7:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    row7 = table.rows[6]

    # Component 2: A7 label is 'Average' (0.2 points)
    try:
        cell_a7 = row7.cells[0].text.strip()
        if cell_a7 == 'Average':
            print(f"PASS: Component 2 — A7 label is 'Average' (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 2 — A7 expected 'Average', found '{cell_a7}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: B7 = '86' (mean of Homework: 85,92,78,95,80) (0.2 points)
    try:
        cell_b7 = row7.cells[1].text.strip()
        # Accept '86' or '86.0' as both represent the mean of 86.0
        b7_numeric = None
        try:
            b7_numeric = float(cell_b7)
        except (ValueError, TypeError):
            pass
        if b7_numeric is not None and abs(b7_numeric - 86.0) < 0.1:
            print(f"PASS: Component 3 — B7 Homework mean is '{cell_b7}' (~86) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — B7 expected '86' (Homework mean), found '{cell_b7}'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: C7 = '82' (mean of Midterm: 78,88,82,90,72) (0.2 points)
    try:
        cell_c7 = row7.cells[2].text.strip()
        c7_numeric = None
        try:
            c7_numeric = float(cell_c7)
        except (ValueError, TypeError):
            pass
        if c7_numeric is not None and abs(c7_numeric - 82.0) < 0.1:
            print(f"PASS: Component 4 — C7 Midterm mean is '{cell_c7}' (~82) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 4 — C7 expected '82' (Midterm mean), found '{cell_c7}'")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: D7 = '86.6' (mean of Final: 90,85,88,92,78) (0.2 points)
    try:
        cell_d7 = row7.cells[3].text.strip()
        d7_numeric = None
        try:
            d7_numeric = float(cell_d7)
        except (ValueError, TypeError):
            pass
        if d7_numeric is not None and abs(d7_numeric - 86.6) < 0.1:
            print(f"PASS: Component 5 — D7 Final mean is '{cell_d7}' (~86.6) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 5 — D7 expected '86.6' (Final mean), found '{cell_d7}'")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/class_grades.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
