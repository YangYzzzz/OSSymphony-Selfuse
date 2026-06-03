"""
Initial Setup: HR Policy Memo with no first-line indentation
Task ID: writer_hr_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_016'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Human Resources Policy Memo', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle / metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Effective Date: January 15, 2026')
    run.font.size = Pt(11)
    run.font.italic = True

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = meta2.add_run('Prepared by: Office of Human Resources')
    run2.font.size = Pt(11)
    run2.font.italic = True

    # --- 8 body paragraphs (NO first-line indent) ---
    body_texts = [
        "This memorandum outlines the revised human resources policies for all full-time "
        "and part-time employees of Meridian Technologies, Inc. These updates reflect changes "
        "in federal and state labor regulations effective for the 2026 fiscal year and supersede "
        "any prior policy documents dated before January 1, 2026.",

        "All employees are expected to maintain a minimum of 40 hours per week for full-time "
        "positions and 20 hours per week for part-time positions. Flex-time arrangements must "
        "be pre-approved by the department manager and submitted through the HR portal at least "
        "two weeks in advance. Unapproved schedule changes may result in disciplinary review.",

        "Paid time off accrues at a rate of 1.5 days per month for employees in their first "
        "three years of service, increasing to 2.0 days per month after the third anniversary. "
        "Unused PTO may be carried over into the following calendar year up to a maximum of 15 "
        "days. Any balance exceeding 15 days will be forfeited unless a written exception is "
        "granted by the VP of Human Resources.",

        "The company maintains a zero-tolerance policy regarding workplace harassment and "
        "discrimination. All incidents must be reported to HR within 48 hours of occurrence. "
        "Investigations will be conducted confidentially, and retaliatory actions against "
        "reporting employees are strictly prohibited under both company policy and applicable law.",

        "Employee performance reviews are conducted on a semi-annual basis, with formal "
        "evaluations scheduled in June and December. Mid-cycle check-ins are encouraged but "
        "not mandatory. Managers are required to complete all evaluation forms using the "
        "standardized template available on the company intranet by the 15th of the review month.",

        "Health insurance enrollment occurs during the annual open enrollment period held each "
        "November. New hires are eligible for coverage beginning on the first day of the month "
        "following 30 days of continuous employment. Qualifying life events such as marriage, "
        "birth of a child, or loss of other coverage permit mid-year enrollment changes.",

        "Professional development opportunities, including tuition reimbursement and conference "
        "attendance, are available to employees who have completed at least one year of service. "
        "Requests must be submitted to the Learning and Development team no fewer than 30 days "
        "prior to the event. The annual cap for tuition reimbursement is $5,250 per employee.",

        "Remote work arrangements are permitted for eligible roles as determined by department "
        "leadership. Employees working remotely must ensure they have a secure internet connection "
        "and a dedicated workspace. Remote workers are subject to the same performance standards "
        "and availability expectations as on-site employees, and must be reachable during core "
        "business hours of 9:00 AM to 3:00 PM in their local time zone.",
    ]

    for text in body_texts:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(6)
        # Explicitly ensure NO first-line indent
        para.paragraph_format.first_line_indent = None

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
