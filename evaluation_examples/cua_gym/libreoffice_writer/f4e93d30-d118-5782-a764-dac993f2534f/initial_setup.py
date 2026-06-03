"""
Initial Setup: Multi-app task - LibreOffice Writer + GIMP image editing
Task ID: osworld_multi_apps_writer_gimp_060
Domain: libreoffice_writer + gimp
Description: Creates client_feedback.docx (feedback document) and product.png (product image
with shadow and off-white background) on the Desktop. Opens client_feedback.docx in
LibreOffice Writer for the agent to read.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFilter
import numpy as np

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_gimp_060'
FEEDBACK_DOC = f'{DESKTOP}/client_feedback.docx'
PRODUCT_IMG = f'{DESKTOP}/product.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_feedback_doc():
    """Create the client feedback document requesting image edits."""
    doc = Document()

    # Title
    title = doc.add_heading('Client Feedback — Product Image Revision Request', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date and reference
    info_para = doc.add_paragraph()
    info_run = info_para.add_run('Date: March 5, 2026    |    Project: Spring Collection 2026    |    Ref: IMG-042')
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # spacer

    # Introduction
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'Dear Design Team,'
    )
    intro_run.font.size = Pt(11)

    doc.add_paragraph()

    body = doc.add_paragraph()
    body_run = body.add_run(
        'We have reviewed the latest product shot for the Spring Collection catalog and have the '
        'following revision requests. Please apply all changes to the file "product.png" and save '
        'the final result as "product_revised.png" on the Desktop.'
    )
    body_run.font.size = Pt(11)

    doc.add_paragraph()

    # Section heading: Requested Changes
    doc.add_heading('Requested Changes', level=2)

    # Change 1 - shadow
    change1_heading = doc.add_paragraph()
    r = change1_heading.add_run('1. Remove Shadow on the Right Side')
    r.bold = True
    r.font.size = Pt(11)

    change1_body = doc.add_paragraph()
    r2 = change1_body.add_run(
        'The current image has a visible drop shadow extending to the right of the product. '
        'This shadow does not match our brand guidelines and must be completely removed. '
        'The right edge of the product should be clean with no shadow artifacts.'
    )
    r2.font.size = Pt(11)
    change1_body.paragraph_format.left_indent = Pt(18)

    doc.add_paragraph()

    # Change 2 - background
    change2_heading = doc.add_paragraph()
    r3 = change2_heading.add_run('2. Change Background to Pure White')
    r3.bold = True
    r3.font.size = Pt(11)

    change2_body = doc.add_paragraph()
    r4 = change2_body.add_run(
        'The background currently appears as an off-white / light grey tone. '
        'Please replace the entire background with pure white (RGB: 255, 255, 255). '
        'This is required for consistent appearance across all print and digital media.'
    )
    r4.font.size = Pt(11)
    change2_body.paragraph_format.left_indent = Pt(18)

    doc.add_paragraph()

    # Change 3 - sharpness
    change3_heading = doc.add_paragraph()
    r5 = change3_heading.add_run('3. Increase Sharpness')
    r5.bold = True
    r5.font.size = Pt(11)

    change3_body = doc.add_paragraph()
    r6 = change3_body.add_run(
        'The product photo appears slightly soft. Please apply a sharpness enhancement to '
        'bring out the fine details of the product texture and edges. '
        'A moderate sharpening pass (factor ~2.0) is recommended.'
    )
    r6.font.size = Pt(11)
    change3_body.paragraph_format.left_indent = Pt(18)

    doc.add_paragraph()

    # Deliverable
    doc.add_heading('Deliverable', level=2)
    deliverable = doc.add_paragraph()
    r7 = deliverable.add_run(
        'Save the revised image as "product_revised.png" on the Desktop. '
        'The final file must show the product on a pure white background, '
        'with the right-side shadow removed, and sharpness enhancement applied.'
    )
    r7.font.size = Pt(11)

    doc.add_paragraph()

    # Sign-off
    signoff = doc.add_paragraph()
    r8 = signoff.add_run('Kind regards,\nSophia Hartmann\nCreative Director, Luminos Retail Group')
    r8.font.size = Pt(11)

    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)
    doc.save(FEEDBACK_DOC)
    print(f'Feedback document created: {FEEDBACK_DOC}')


def create_product_image():
    """Create a realistic product image with a shadow on the right and off-white background."""
    width, height = 800, 600

    # Off-white / light grey background (NOT pure white — agent must fix this)
    bg_color = (235, 233, 230)
    img = Image.new('RGB', (width, height), color=bg_color)
    draw = ImageDraw.Draw(img)

    # Draw a product: a stylized bottle/box shape (product silhouette)
    # Product body (dark blue/navy — simulates a product bottle)
    product_x1, product_y1 = 280, 100
    product_x2, product_y2 = 480, 480

    # Main product body
    draw.rectangle([product_x1, product_y1, product_x2, product_y2],
                   fill=(30, 55, 100), outline=(20, 40, 80), width=2)

    # Label area on product
    label_x1, label_y1 = 295, 200
    label_x2, label_y2 = 465, 380
    draw.rectangle([label_x1, label_y1, label_x2, label_y2],
                   fill=(240, 240, 245), outline=(200, 200, 210), width=1)

    # Product cap
    cap_x1, cap_y1 = 310, 60
    cap_x2, cap_y2 = 450, 105
    draw.rectangle([cap_x1, cap_y1, cap_x2, cap_y2],
                   fill=(180, 160, 30), outline=(150, 130, 20), width=2)

    # Label text lines (simulated)
    draw.rectangle([310, 220, 450, 235], fill=(30, 55, 100))
    draw.rectangle([310, 250, 440, 260], fill=(150, 150, 160))
    draw.rectangle([310, 270, 435, 278], fill=(150, 150, 160))
    draw.rectangle([310, 290, 420, 298], fill=(150, 150, 160))
    draw.rectangle([310, 330, 440, 345], fill=(30, 55, 100))
    draw.rectangle([310, 355, 430, 362], fill=(150, 150, 160))

    # Shadow on the RIGHT side of the product (gradient-like, going right)
    # This is what the agent needs to remove
    shadow_img = img.copy()
    shadow_arr = np.array(shadow_img, dtype=np.float32)

    # Create a shadow gradient on the right side of the product
    shadow_x_start = product_x2
    shadow_x_end = min(product_x2 + 100, width)
    shadow_y_start = product_y1 + 30
    shadow_y_end = product_y2 - 30

    for x in range(shadow_x_start, shadow_x_end):
        alpha_factor = 1.0 - (x - shadow_x_start) / (shadow_x_end - shadow_x_start)
        intensity = int(60 * alpha_factor)  # shadow darkening
        for y in range(shadow_y_start, shadow_y_end):
            shadow_arr[y, x, 0] = max(0, shadow_arr[y, x, 0] - intensity)
            shadow_arr[y, x, 1] = max(0, shadow_arr[y, x, 1] - intensity)
            shadow_arr[y, x, 2] = max(0, shadow_arr[y, x, 2] - intensity)

    img = Image.fromarray(shadow_arr.astype(np.uint8))

    # Apply a very slight blur to the shadow edge to make it more realistic
    # Only blur the shadow region
    shadow_region = img.crop((shadow_x_start, shadow_y_start, shadow_x_end, shadow_y_end))
    shadow_region_blurred = shadow_region.filter(ImageFilter.GaussianBlur(radius=3))
    img.paste(shadow_region_blurred, (shadow_x_start, shadow_y_start))

    os.makedirs(DESKTOP, exist_ok=True)
    img.save(PRODUCT_IMG)
    print(f'Product image created: {PRODUCT_IMG}')


def create_initial():
    create_feedback_doc()
    create_product_image()

    # GUI-ready startup: open client_feedback.docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{FEEDBACK_DOC}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with client_feedback.docx (DISPLAY=:0)')


create_initial()
