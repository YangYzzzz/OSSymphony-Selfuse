"""
Initial Setup: Writer document with three heading levels, no automatic numbering.
Task ID: writer_bs_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default style font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('Quarterly Performance Review Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # === SECTION 1: Introduction (Heading 1) ===
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'This report presents a comprehensive analysis of the quarterly '
        'performance metrics across all departments within Meridian Technologies. '
        'The review period covers January through March 2025 and includes data '
        'from twelve regional offices.'
    )

    # -- 1. Background (Heading 2) --
    doc.add_heading('Background', level=2)

    doc.add_paragraph(
        'Meridian Technologies has experienced significant growth over the past '
        'fiscal year, expanding into three new markets including Southeast Asia, '
        'Eastern Europe, and South America. This expansion has necessitated a '
        'thorough review of operational efficiency and resource allocation.'
    )

    # -- 1.1. Overview (Heading 3) --
    doc.add_heading('Overview', level=3)

    doc.add_paragraph(
        'The performance evaluation framework was redesigned in Q4 2024 to '
        'incorporate both quantitative metrics and qualitative assessments. Key '
        'performance indicators now include customer satisfaction scores, revenue '
        'per employee, and project delivery timelines.'
    )

    # -- 1.2. Scope (Heading 3) --
    doc.add_heading('Scope', level=3)

    doc.add_paragraph(
        'This review covers all departments including Engineering, Sales, '
        'Marketing, Human Resources, and Operations. Data was collected from '
        'the internal analytics platform and validated by department leads.'
    )

    # -- 2. Objectives (Heading 2) --
    doc.add_heading('Objectives', level=2)

    doc.add_paragraph(
        'The primary objectives of this quarterly review are to identify areas '
        'of improvement, recognize high-performing teams, and establish '
        'benchmarks for the upcoming quarter. Additionally, the review aims to '
        'align departmental goals with the company-wide strategic vision.'
    )

    # === SECTION 2: Methods (Heading 1) ===
    doc.add_heading('Methods', level=1)

    doc.add_paragraph(
        'The methodology employed for this performance review combines automated '
        'data collection with structured interviews and peer assessments. Each '
        'department was evaluated using a standardized rubric developed by the '
        'Human Resources team in collaboration with senior leadership.'
    )

    # -- 1. Data Collection (Heading 2) --
    doc.add_heading('Data Collection', level=2)

    doc.add_paragraph(
        'Performance data was gathered from multiple sources including the '
        'enterprise resource planning system, customer relationship management '
        'platform, and project management tools. A total of 2,847 data points '
        'were collected across all departments.'
    )

    # -- 1.1. Survey Design (Heading 3) --
    doc.add_heading('Survey Design', level=3)

    doc.add_paragraph(
        'Employee satisfaction surveys were distributed electronically to all '
        '1,243 full-time employees. The survey instrument consisted of 35 '
        'questions using a five-point Likert scale, with an overall response '
        'rate of 78.4 percent.'
    )

    # -- 1.2. Interview Protocol (Heading 3) --
    doc.add_heading('Interview Protocol', level=3)

    doc.add_paragraph(
        'Semi-structured interviews were conducted with 48 team leads and '
        'department managers. Each interview lasted approximately 45 minutes '
        'and covered topics including team dynamics, resource constraints, '
        'and strategic alignment.'
    )

    # -- 2. Statistical Analysis (Heading 2) --
    doc.add_heading('Statistical Analysis', level=2)

    doc.add_paragraph(
        'Quantitative data was analyzed using descriptive statistics and '
        'trend analysis. Year-over-year comparisons were performed to identify '
        'significant changes in key performance indicators. Statistical '
        'significance was set at p < 0.05 for all tests.'
    )

    # === SECTION 3: Results (Heading 1) ===
    doc.add_heading('Results', level=1)

    doc.add_paragraph(
        'The results of the quarterly review indicate overall positive trends '
        'across most departments, with notable improvements in customer '
        'satisfaction and project delivery timelines. Revenue targets were met '
        'or exceeded by eight of the twelve regional offices.'
    )

    # -- 1. Financial Performance (Heading 2) --
    doc.add_heading('Financial Performance', level=2)

    doc.add_paragraph(
        'Total revenue for Q1 2025 reached $47.3 million, representing a '
        '12.8 percent increase compared to Q1 2024. The Engineering department '
        'contributed $18.9 million through product licensing and custom '
        'development services.'
    )

    # -- 1.1. Revenue Breakdown (Heading 3) --
    doc.add_heading('Revenue Breakdown', level=3)

    doc.add_paragraph(
        'Product licensing accounted for 62 percent of total revenue, followed '
        'by professional services at 24 percent and support contracts at 14 '
        'percent. The Southeast Asia market generated $3.2 million in its first '
        'quarter of operation.'
    )

    # -- 2. Employee Metrics (Heading 2) --
    doc.add_heading('Employee Metrics', level=2)

    doc.add_paragraph(
        'Employee retention improved to 94.2 percent, up from 91.7 percent in '
        'the previous quarter. Average employee satisfaction score increased '
        'from 3.8 to 4.1 on the five-point scale. Training completion rates '
        'reached 88 percent across all departments.'
    )

    # -- 2.1. Productivity Analysis (Heading 3) --
    doc.add_heading('Productivity Analysis', level=3)

    doc.add_paragraph(
        'Average project delivery time decreased by 15 percent compared to '
        'the previous quarter. The Engineering team completed 94 percent of '
        'sprint goals, while the Marketing team achieved a 23 percent increase '
        'in campaign conversion rates.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
