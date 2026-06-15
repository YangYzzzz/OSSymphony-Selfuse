"""
FINAL REWARD SCRIPT - SUCCESS
Task: I’m re-formatting my table and need to break one cell apart. In Table 1, could you walk me through turning the single cell at C2 into exactly 2 separate rows?
Generated: 2025-09-10 15:47:23
Status: success
Model: azure-o3
Total Steps: 15
"""

import os
from docx import Document
from docx.oxml.ns import qn

def verify_split_cell(file_path: str) -> float:
    """Verify that the original merged cell C2 of Table 1 has been split into
    exactly two independent rows.

    Progressive scoring (total 1.0):
      • 0.6 points – Cells (1,2) and (2,2) are DIFFERENT objects and neither
        contains a <w:vMerge> tag (i.e., no vertical merge).
      • 0.4 points – The two rows created by the split have identical values
        in columns A and B, indicating that only C2 was split while the rest
        of the row data was duplicated (heuristic but strong evidence of a
        proper split rather than arbitrary edits).

    The function prints detailed diagnostics and always returns a float in the
    range [0.0, 1.0].
    """
    print(f"Starting verification for: {file_path}")

    max_score = 1.0
    score = 0.0

    # ---- Basic file checks (no points for these) ----
    if not os.path.exists(file_path):
        print("✗ Document not found")
        return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to load DOCX: {e}")
        return 0.0
    if not doc.tables:
        print("✗ No tables found in document")
        return 0.0

    tbl = doc.tables[0]  # Task refers to “Table 1” – the first table
    rows, cols = len(tbl.rows), len(tbl.columns)
    print(f"Found first table with {rows} rows and {cols} columns (pre-check – no points)")

    # ---- REQUIREMENT 1: C2 is now two independent cells (0.6) ----
    if rows >= 3 and cols >= 3:
        top_cell   = tbl.cell(1, 2)  # row index 1 (2nd row), col index 2 (column C)
        bottom_cell = tbl.cell(2, 2)  # row index 2 (3rd row)

        same_object = top_cell._tc is bottom_cell._tc
        print(f"  Debug: cell objects are {'the SAME' if same_object else 'DIFFERENT'}")

        def has_vmerge(cell):
            return any(True for _ in cell._tc.iter(qn('w:vMerge')))

        top_vmerge    = has_vmerge(top_cell)
        bottom_vmerge = has_vmerge(bottom_cell)
        print(f"  Debug: top cell has vMerge: {top_vmerge}")
        print(f"  Debug: bottom cell has vMerge: {bottom_vmerge}")

        if (not same_object) and (not top_vmerge) and (not bottom_vmerge):
            print("✓ Requirement 1 passed: cells are independent and unmerged (0.6)")
            score += 0.6
        else:
            print("✗ Requirement 1 failed: cells appear merged or share <w:vMerge> tags")
    else:
        print("✗ Table too small to evaluate Requirement 1")

    # ---- REQUIREMENT 2: Duplicate row data in A & B evidences a split (0.4) ----
    if rows >= 3 and cols >= 2:
        text_A_top    = tbl.cell(1, 0).text.strip()
        text_A_bottom = tbl.cell(2, 0).text.strip()
        text_B_top    = tbl.cell(1, 1).text.strip()
        text_B_bottom = tbl.cell(2, 1).text.strip()

        duplicate_rows = (text_A_top == text_A_bottom) and (text_B_top == text_B_bottom)
        print(f"  Debug: Columns A & B duplicated between the split rows: {duplicate_rows}")

        if duplicate_rows:
            print("✓ Requirement 2 passed: duplicated row data consistent with a split (0.4)")
            score += 0.4
        else:
            print("✗ Requirement 2 failed: row data not duplicated as expected")
    else:
        print("✗ Table too small to evaluate Requirement 2")

    # ---- Final score ----
    final_score = min(score, max_score)
    print(f"Total Score = {final_score}/{max_score}")
    return final_score


if __name__ == "__main__":
    test_path = "/home/user/im_re_formatting_my_table_and_need_to_break_one_cell_apart_in_table_1_could_you_walk_me_through_turn.docx"
    reward = verify_split_cell(test_path)
    print(f"REWARD: {reward}")
