"""
Reward Script: Center-align, bold, and resize title in legal document
Task ID: writer_legal_014
Domain: libreoffice_writer
Scoring:
  Component 1: Title paragraph is center-aligned (0.35 pts)
  Component 2: Title text is bold (0.35 pts)
  Component 3: Title font size is 14pt (0.30 pts)
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_014'
EXPECTED_TITLE = 'SETTLEMENT AGREEMENT AND MUTUAL RELEASE'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: first paragraph must contain the expected title text
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    para = doc.paragraphs[0]
    if para.text.strip() != EXPECTED_TITLE:
        print(f"FAIL: First paragraph text mismatch. Expected: '{EXPECTED_TITLE}', Found: '{para.text.strip()}'")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title paragraph is center-aligned (0.35 points)
    try:
        alignment = para.paragraph_format.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            print(f"PASS: Component 1 — Title is center-aligned (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — Expected CENTER alignment, found: {alignment}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Title text is bold (0.35 points)
    # Check that all runs in the title paragraph are bold
    try:
        runs = para.runs
        if len(runs) == 0:
            print("FAIL: Component 2 — No runs found in title paragraph")
        else:
            all_bold = all(run.font.bold is True for run in runs if run.text.strip())
            if all_bold:
                print(f"PASS: Component 2 — Title is bold (0.35 pts)")
                total_score += 0.35
            else:
                bold_states = [(run.text[:20], run.font.bold) for run in runs]
                print(f"FAIL: Component 2 — Not all runs are bold. States: {bold_states}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Title font size is 14pt (0.30 points)
    # Check that all runs in the title paragraph have 14pt font size
    try:
        runs = para.runs
        if len(runs) == 0:
            print("FAIL: Component 3 — No runs found in title paragraph")
        else:
            all_14pt = all(
                run.font.size is not None and abs(run.font.size.pt - 14.0) < 0.5
                for run in runs if run.text.strip()
            )
            if all_14pt:
                sizes = [run.font.size.pt for run in runs if run.text.strip() and run.font.size]
                print(f"PASS: Component 3 — Title font size is 14pt (actual: {sizes}) (0.30 pts)")
                total_score += 0.30
            else:
                sizes = [(run.text[:20], run.font.size.pt if run.font.size else None) for run in runs]
                print(f"FAIL: Component 3 — Expected 14pt, found: {sizes}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
