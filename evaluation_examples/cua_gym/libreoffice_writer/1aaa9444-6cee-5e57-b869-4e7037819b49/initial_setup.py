"""
Initial Setup: Legal settlement agreement document with left-aligned 12pt regular title
Task ID: writer_legal_014
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_014'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title - LEFT aligned, 12pt, NOT bold (task requires agent to center + bold + 14pt)
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run('SETTLEMENT AGREEMENT AND MUTUAL RELEASE')
    title_run.font.size = Pt(12)
    title_run.font.bold = False
    title_run.font.name = 'Times New Roman'

    # Preamble
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(6)
    r = preamble.add_run(
        'This Settlement Agreement and Mutual Release ("Agreement") is entered into '
        'as of March 15, 2025, by and between:'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Party A
    party_a = doc.add_paragraph()
    party_a.paragraph_format.space_after = Pt(6)
    r = party_a.add_run(
        'Meridian Technologies, Inc., a Delaware corporation with its principal place of '
        'business at 4200 Innovation Drive, Suite 800, San Jose, California 95134 '
        '("Meridian" or "Party A");'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # And
    and_para = doc.add_paragraph()
    and_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    and_para.paragraph_format.space_after = Pt(6)
    r = and_para.add_run('and')
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Party B
    party_b = doc.add_paragraph()
    party_b.paragraph_format.space_after = Pt(6)
    r = party_b.add_run(
        'Crestview Solutions, LLC, a California limited liability company with its '
        'principal place of business at 1875 Wilshire Boulevard, Suite 1200, '
        'Los Angeles, California 90036 ("Crestview" or "Party B").'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Collectively
    collective = doc.add_paragraph()
    collective.paragraph_format.space_after = Pt(12)
    r = collective.add_run(
        'Meridian and Crestview are collectively referred to herein as the "Parties" '
        'and individually as a "Party."'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Recitals heading
    recitals_heading = doc.add_paragraph()
    recitals_heading.paragraph_format.space_before = Pt(12)
    recitals_heading.paragraph_format.space_after = Pt(6)
    r = recitals_heading.add_run('RECITALS')
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.name = 'Times New Roman'

    # Recital A
    rec_a = doc.add_paragraph()
    rec_a.paragraph_format.space_after = Pt(6)
    r = rec_a.add_run(
        'A. WHEREAS, Meridian and Crestview entered into a Software Development and '
        'Licensing Agreement dated June 1, 2023 (the "Original Agreement") pursuant to '
        'which Crestview agreed to develop certain proprietary software modules for '
        "Meridian's enterprise resource planning platform;"
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Recital B
    rec_b = doc.add_paragraph()
    rec_b.paragraph_format.space_after = Pt(6)
    r = rec_b.add_run(
        'B. WHEREAS, disputes have arisen between the Parties regarding the scope of '
        'deliverables, milestone timelines, and payment obligations under the Original '
        'Agreement, which disputes led to the filing of Case No. 2024-CV-08731 in the '
        'Superior Court of California, County of Santa Clara (the "Litigation");'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Recital C
    rec_c = doc.add_paragraph()
    rec_c.paragraph_format.space_after = Pt(6)
    r = rec_c.add_run(
        'C. WHEREAS, the Parties desire to resolve all disputes and claims between them '
        'arising from or related to the Original Agreement and the Litigation, without '
        'any admission of liability by either Party, and to enter into this Agreement on '
        'the terms and conditions set forth herein;'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Agreement heading
    agreement_heading = doc.add_paragraph()
    agreement_heading.paragraph_format.space_before = Pt(12)
    agreement_heading.paragraph_format.space_after = Pt(6)
    r = agreement_heading.add_run('AGREEMENT')
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.name = 'Times New Roman'

    # Section 1
    sec1 = doc.add_paragraph()
    sec1.paragraph_format.space_after = Pt(6)
    r = sec1.add_run(
        '1. Settlement Payment. In consideration of the mutual promises and covenants '
        'contained herein, Meridian shall pay to Crestview the total sum of Seven '
        'Hundred Fifty Thousand Dollars ($750,000.00) (the "Settlement Amount"), '
        'payable as follows:'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Sub-section 1a
    sec1a = doc.add_paragraph()
    sec1a.paragraph_format.left_indent = Inches(0.5)
    sec1a.paragraph_format.space_after = Pt(6)
    r = sec1a.add_run(
        '(a) Three Hundred Seventy-Five Thousand Dollars ($375,000.00) within thirty '
        '(30) calendar days of the Effective Date; and'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Sub-section 1b
    sec1b = doc.add_paragraph()
    sec1b.paragraph_format.left_indent = Inches(0.5)
    sec1b.paragraph_format.space_after = Pt(6)
    r = sec1b.add_run(
        '(b) Three Hundred Seventy-Five Thousand Dollars ($375,000.00) within ninety '
        '(90) calendar days of the Effective Date.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Section 2
    sec2 = doc.add_paragraph()
    sec2.paragraph_format.space_after = Pt(6)
    r = sec2.add_run(
        '2. Mutual Release. Upon receipt of the Settlement Amount in full, each Party, '
        'on behalf of itself and its officers, directors, employees, agents, successors, '
        'and assigns, hereby releases and forever discharges the other Party from any '
        'and all claims, demands, damages, actions, causes of action, suits, costs, '
        'expenses, and liabilities of any kind or nature, whether known or unknown, '
        'arising out of or related to the Original Agreement or the Litigation.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Section 3
    sec3 = doc.add_paragraph()
    sec3.paragraph_format.space_after = Pt(6)
    r = sec3.add_run(
        '3. Confidentiality. The Parties agree that the terms of this Agreement, '
        'including the Settlement Amount, shall be kept strictly confidential and shall '
        'not be disclosed to any third party, except as required by law, regulation, or '
        'court order, or to the Parties\' respective legal counsel, accountants, and '
        'tax advisors.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Section 4
    sec4 = doc.add_paragraph()
    sec4.paragraph_format.space_after = Pt(6)
    r = sec4.add_run(
        '4. Governing Law. This Agreement shall be governed by and construed in '
        'accordance with the laws of the State of California, without regard to its '
        'conflicts of law principles.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
