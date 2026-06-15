"""
Initial Setup: Requirements table with no alternating row colors
Task ID: writer_tech_042
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_042'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Software Requirements Specification', level=1)

    # Introductory paragraph
    doc.add_paragraph(
        'This document outlines the functional and non-functional requirements '
        'for the CloudSync Platform v3.2 release. All requirements have been '
        'reviewed and prioritized by the product management team.'
    )

    doc.add_paragraph(
        'The following table summarizes the key requirements, their priority '
        'levels, current status, and assigned teams.'
    )

    # Create the requirements table: header row + 10 data rows = 11 rows total
    table = doc.add_table(rows=11, cols=5)
    table.style = 'Table Grid'

    # Header row
    headers = ['Req ID', 'Description', 'Priority', 'Status', 'Assigned Team']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # 10 rows of realistic requirements data
    data = [
        ['REQ-001', 'Implement OAuth 2.0 single sign-on with SAML fallback', 'Critical', 'In Progress', 'Auth Platform'],
        ['REQ-002', 'Add real-time file sync with conflict resolution for shared folders', 'High', 'Open', 'Sync Engine'],
        ['REQ-003', 'Support bulk export of user data in CSV and JSON formats', 'Medium', 'Completed', 'Data Services'],
        ['REQ-004', 'Integrate Stripe payment gateway for subscription billing', 'Critical', 'In Progress', 'Billing'],
        ['REQ-005', 'Implement role-based access control with custom permission sets', 'High', 'Open', 'Auth Platform'],
        ['REQ-006', 'Add automated backup scheduling with retention policies', 'Medium', 'In Review', 'Infrastructure'],
        ['REQ-007', 'Build dashboard with real-time analytics and usage metrics', 'High', 'In Progress', 'Analytics'],
        ['REQ-008', 'Enable two-factor authentication via TOTP and SMS', 'Critical', 'Completed', 'Auth Platform'],
        ['REQ-009', 'Create REST API rate limiting with configurable thresholds', 'Medium', 'Open', 'API Gateway'],
        ['REQ-010', 'Add audit logging for all administrative actions', 'High', 'In Review', 'Compliance'],
    ]

    for i, row_data in enumerate(data, 1):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    # Set column widths for readability
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(3.5)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(1.0)
        row.cells[4].width = Inches(1.2)

    # Additional context paragraph after the table
    doc.add_paragraph('')
    doc.add_paragraph(
        'Note: Requirements marked as Critical must be completed before the '
        'v3.2 release milestone on June 30, 2025. All In Review items require '
        'sign-off from the technical lead before moving to development.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
