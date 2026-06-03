"""
Initial Setup: Legal agreement document with plain-text article and section headings.
Task ID: writer_lec_008
Domain: libreoffice_writer

Creates a legal agreement document where article headings (e.g., "Article 1 - Definitions")
and section headings (e.g., "Section 1.1 - Key Terms") are plain text paragraphs
with no list numbering styles applied.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_article_heading(doc, text):
    """Add an article-level heading as plain bold paragraph (no list style)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_section_heading(doc, text):
    """Add a section-level heading as plain bold-italic paragraph (no list style)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.25)
    return para


def add_body_text(doc, text):
    """Add a normal body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Inches(0.5)
    return para


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run('MASTER SERVICES AGREEMENT')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'

    # --- Preamble ---
    preamble = doc.add_paragraph()
    run = preamble.add_run(
        'This Master Services Agreement ("Agreement") is entered into as of March 15, 2025, '
        'by and between Meridian Technologies Inc., a Delaware corporation ("Provider"), '
        'and Cascade Financial Group LLC, an Oregon limited liability company ("Client"). '
        'Both parties agree to the following terms and conditions:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    preamble.paragraph_format.space_after = Pt(12)

    # --- Article 1: Definitions ---
    add_article_heading(doc, 'Definitions')

    add_section_heading(doc, 'Key Terms')
    add_body_text(doc,
        '"Confidential Information" means any proprietary data, trade secrets, technical '
        'specifications, business strategies, customer lists, or financial records disclosed '
        'by either party during the term of this Agreement.'
    )

    add_section_heading(doc, 'Service Scope')
    add_body_text(doc,
        '"Services" refers to the cloud infrastructure management, cybersecurity auditing, '
        'and data analytics consulting described in Exhibit A attached hereto.'
    )

    add_section_heading(doc, 'Deliverables')
    add_body_text(doc,
        '"Deliverables" means all reports, software modules, documentation, and tangible '
        'work product created by Provider in fulfillment of the Services.'
    )

    # --- Article 2: Term and Termination ---
    add_article_heading(doc, 'Term and Termination')

    add_section_heading(doc, 'Effective Date')
    add_body_text(doc,
        'This Agreement shall commence on the date first written above and shall continue '
        'for a period of twenty-four (24) months unless terminated earlier in accordance '
        'with the provisions herein.'
    )

    add_section_heading(doc, 'Termination for Convenience')
    add_body_text(doc,
        'Either party may terminate this Agreement upon sixty (60) days prior written notice '
        'to the other party. Such termination shall not relieve the Client of payment '
        'obligations for Services already rendered.'
    )

    add_section_heading(doc, 'Termination for Cause')
    add_body_text(doc,
        'Either party may terminate this Agreement immediately upon written notice if the '
        'other party materially breaches any obligation hereunder and fails to cure such '
        'breach within thirty (30) days of receiving written notice thereof.'
    )

    # --- Article 3: Compensation and Payment ---
    add_article_heading(doc, 'Compensation and Payment')

    add_section_heading(doc, 'Fee Schedule')
    add_body_text(doc,
        'Client shall pay Provider a monthly retainer of $45,000 for ongoing infrastructure '
        'management services, plus $275 per hour for additional consulting engagements as '
        'outlined in the Statement of Work.'
    )

    add_section_heading(doc, 'Payment Terms')
    add_body_text(doc,
        'All invoices are due within thirty (30) days of receipt. Late payments shall accrue '
        'interest at a rate of 1.5% per month or the maximum rate permitted by law, '
        'whichever is less.'
    )

    # --- Article 4: Intellectual Property ---
    add_article_heading(doc, 'Intellectual Property')

    add_section_heading(doc, 'Ownership of Deliverables')
    add_body_text(doc,
        'Upon full payment, all Deliverables created specifically for Client under this '
        'Agreement shall become the exclusive property of Client. Provider retains ownership '
        'of all pre-existing tools, frameworks, and methodologies.'
    )

    add_section_heading(doc, 'License Grant')
    add_body_text(doc,
        'Provider grants Client a non-exclusive, perpetual, royalty-free license to use any '
        'pre-existing Provider intellectual property embedded in the Deliverables, solely '
        'for Client\'s internal business purposes.'
    )

    # --- Article 5: Confidentiality ---
    add_article_heading(doc, 'Confidentiality')

    add_section_heading(doc, 'Non-Disclosure Obligations')
    add_body_text(doc,
        'Each party agrees to hold the other party\'s Confidential Information in strict '
        'confidence and not to disclose it to any third party without prior written consent, '
        'except as required by law or regulation.'
    )

    add_section_heading(doc, 'Duration of Obligations')
    add_body_text(doc,
        'The obligations of confidentiality shall survive the termination or expiration of '
        'this Agreement for a period of five (5) years from the date of disclosure.'
    )

    # --- Article 6: Liability and Indemnification ---
    add_article_heading(doc, 'Liability and Indemnification')

    add_section_heading(doc, 'Limitation of Liability')
    add_body_text(doc,
        'In no event shall either party be liable for indirect, incidental, consequential, '
        'special, or punitive damages arising out of this Agreement, regardless of the form '
        'of action or theory of liability, even if advised of the possibility of such damages.'
    )

    add_section_heading(doc, 'Indemnification')
    add_body_text(doc,
        'Provider shall indemnify, defend, and hold harmless Client from and against any '
        'claims, damages, losses, or expenses arising from Provider\'s negligent acts or '
        'omissions in the performance of Services under this Agreement.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
