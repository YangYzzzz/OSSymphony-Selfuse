"""
Reward Script: Create a custom 'Legal Numbering' list style with Article/Section formatting
Task ID: writer_lec_008
Domain: libreoffice_writer
Scoring:
  Component 1: 'Legal Numbering' list style exists (0.2 pts)
  Component 2: Level 0 uses 'Article' prefix with decimal numbering (0.2 pts)
  Component 3: Level 1 uses 'Section' prefix with parent-child numbering (0.2 pts)
  Component 4: Article headings (ilvl=0) have numbering applied (0.2 pts)
  Component 5: Section headings (ilvl=1) have numbering applied (0.2 pts)
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_008'


def persist_app_state(domain):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def get_abstract_num_for_style(doc, style_name):
    """
    Given a list style name, find the abstractNum definition it references.
    Returns the abstractNum XML element or None.
    """
    # Find the style element
    target_num_id = None
    for s in doc.styles:
        if s.name == style_name:
            numPr = s.element.find(qn('w:pPr') + '/' + qn('w:numPr'))
            if numPr is not None:
                numId_el = numPr.find(qn('w:numId'))
                if numId_el is not None:
                    target_num_id = numId_el.get(qn('w:val'))
            break

    if target_num_id is None:
        return None, None

    # Resolve numId -> abstractNumId
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return None, None

    numbering_xml = numbering_part._element
    abs_num_id = None
    for num in numbering_xml.findall(qn('w:num')):
        if num.get(qn('w:numId')) == target_num_id:
            absNumId_el = num.find(qn('w:abstractNumId'))
            if absNumId_el is not None:
                abs_num_id = absNumId_el.get(qn('w:val'))
            break

    if abs_num_id is None:
        return None, target_num_id

    # Find the abstractNum definition
    for absNum in numbering_xml.findall(qn('w:abstractNum')):
        if absNum.get(qn('w:abstractNumId')) == abs_num_id:
            return absNum, target_num_id

    return None, target_num_id


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 'Legal Numbering' list style exists (0.2 points)
    try:
        legal_style_found = False
        for s in doc.styles:
            if s.name == 'Legal Numbering':
                # Verify it is a numbering/list type style
                style_type = str(s.type)
                if 'LIST' in style_type or 'NUMBERING' in style_type or s.element.get(qn('w:type')) == 'numbering':
                    legal_style_found = True
                    print(f"PASS: Component 1 - 'Legal Numbering' list style exists (type={style_type}) (0.2 pts)")
                    total_score += 0.2
                else:
                    print(f"FAIL: Component 1 - Found 'Legal Numbering' but wrong type: {style_type}")
                break
        if not legal_style_found and total_score == 0.0:
            # Also check for case-insensitive match
            for s in doc.styles:
                if s.name.lower().replace(' ', '') == 'legalnumbering':
                    style_type = str(s.type)
                    if 'LIST' in style_type or 'NUMBERING' in style_type or s.element.get(qn('w:type')) == 'numbering':
                        legal_style_found = True
                        print(f"PASS: Component 1 - 'Legal Numbering' list style exists (case-insensitive match, name={s.name}) (0.2 pts)")
                        total_score += 0.2
                    break
            if not legal_style_found:
                print(f"FAIL: Component 1 - 'Legal Numbering' style not found")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Level 0 uses 'Article' prefix with decimal numbering (0.2 points)
    try:
        absNum, numId = get_abstract_num_for_style(doc, 'Legal Numbering')
        if absNum is None:
            # Try to find numbering from paragraphs directly
            # Check any paragraph that has numbering and ilvl=0
            print(f"FAIL: Component 2 - Could not resolve abstract numbering for 'Legal Numbering' style")
        else:
            lvl0_found = False
            for lvl in absNum.findall(qn('w:lvl')):
                if lvl.get(qn('w:ilvl')) == '0':
                    numFmt_el = lvl.find(qn('w:numFmt'))
                    lvlText_el = lvl.find(qn('w:lvlText'))
                    numFmt = numFmt_el.get(qn('w:val')) if numFmt_el is not None else 'None'
                    lvlText = lvlText_el.get(qn('w:val')) if lvlText_el is not None else 'None'

                    # Check: format should be "Article %1" with decimal numbering
                    # Accept variations like "Article %1", "Article %1.", "Article  %1"
                    article_ok = ('article' in lvlText.lower() and '%1' in lvlText)
                    decimal_ok = (numFmt == 'decimal')

                    if article_ok and decimal_ok:
                        print(f"PASS: Component 2 - Level 0 format: '{lvlText}' with numFmt={numFmt} (0.2 pts)")
                        total_score += 0.2
                        lvl0_found = True
                    else:
                        print(f"FAIL: Component 2 - Level 0: lvlText='{lvlText}', numFmt={numFmt}. Expected 'Article %1' with decimal.")
                    break
            if not lvl0_found and total_score < 0.4:
                print(f"FAIL: Component 2 - Level 0 definition not found in abstract numbering")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Level 1 uses 'Section' prefix with parent-child numbering (0.2 points)
    try:
        if absNum is None:
            print(f"FAIL: Component 3 - No abstract numbering resolved")
        else:
            lvl1_found = False
            for lvl in absNum.findall(qn('w:lvl')):
                if lvl.get(qn('w:ilvl')) == '1':
                    numFmt_el = lvl.find(qn('w:numFmt'))
                    lvlText_el = lvl.find(qn('w:lvlText'))
                    numFmt = numFmt_el.get(qn('w:val')) if numFmt_el is not None else 'None'
                    lvlText = lvlText_el.get(qn('w:val')) if lvlText_el is not None else 'None'

                    # Check: format should be "Section %1.%2" with decimal numbering
                    # This gives "Section 1.1", "Section 1.2", etc.
                    section_ok = ('section' in lvlText.lower() and '%1' in lvlText and '%2' in lvlText)
                    decimal_ok = (numFmt == 'decimal')

                    if section_ok and decimal_ok:
                        print(f"PASS: Component 3 - Level 1 format: '{lvlText}' with numFmt={numFmt} (0.2 pts)")
                        total_score += 0.2
                        lvl1_found = True
                    else:
                        print(f"FAIL: Component 3 - Level 1: lvlText='{lvlText}', numFmt={numFmt}. Expected 'Section %1.%2' with decimal.")
                    break
            if not lvl1_found and total_score < 0.6:
                print(f"FAIL: Component 3 - Level 1 definition not found in abstract numbering")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Article headings (ilvl=0) have numbering applied (0.2 points)
    # The article headings are the major sections: Definitions, Term and Termination,
    # Compensation and Payment, Intellectual Property, Confidentiality, Liability and Indemnification
    try:
        article_paras = []
        for i, para in enumerate(doc.paragraphs):
            numPr = para._element.find(qn('w:pPr') + '/' + qn('w:numPr'))
            if numPr is not None:
                ilvl_el = numPr.find(qn('w:ilvl'))
                if ilvl_el is not None and ilvl_el.get(qn('w:val')) == '0':
                    article_paras.append(para.text.strip())

        # Expect at least 4 article-level paragraphs with numbering
        if len(article_paras) >= 4:
            print(f"PASS: Component 4 - {len(article_paras)} article headings (ilvl=0) have numbering: {article_paras[:3]}... (0.2 pts)")
            total_score += 0.2
        elif len(article_paras) > 0:
            partial = 0.2 * (len(article_paras) / 6.0)
            print(f"PARTIAL: Component 4 - Only {len(article_paras)} article headings have numbering (expected >= 4): {article_paras} ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - No paragraphs with ilvl=0 numbering found")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Section headings (ilvl=1) have numbering applied (0.2 points)
    # Section headings are sub-items under each article
    try:
        section_paras = []
        for i, para in enumerate(doc.paragraphs):
            numPr = para._element.find(qn('w:pPr') + '/' + qn('w:numPr'))
            if numPr is not None:
                ilvl_el = numPr.find(qn('w:ilvl'))
                if ilvl_el is not None and ilvl_el.get(qn('w:val')) == '1':
                    section_paras.append(para.text.strip())

        # Expect at least 6 section-level paragraphs with numbering
        if len(section_paras) >= 6:
            print(f"PASS: Component 5 - {len(section_paras)} section headings (ilvl=1) have numbering: {section_paras[:3]}... (0.2 pts)")
            total_score += 0.2
        elif len(section_paras) > 0:
            partial = 0.2 * (len(section_paras) / 14.0)
            print(f"PARTIAL: Component 5 - Only {len(section_paras)} section headings have numbering (expected >= 6): {section_paras[:3]}... ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 - No paragraphs with ilvl=1 numbering found")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
