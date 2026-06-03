"""
Initial Setup: Book manuscript with 5 chapters, all using Default Page Style
Task ID: writer_rd_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default page margins to 2.54 cm (1 inch) for the single section
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Chapter 1: The Awakening ---
    h1 = doc.add_heading('Chapter 1: The Awakening', level=1)
    doc.add_paragraph(
        'The morning sun crept through the curtains of Elena Vasquez\'s apartment on Maple Street, '
        'casting long golden shadows across the hardwood floor. She had barely slept, her mind racing '
        'with the events of the previous evening. The letter from her grandmother\'s estate attorney '
        'had arrived without warning, its contents both bewildering and intriguing.'
    )
    doc.add_paragraph(
        'Elena poured herself a cup of dark roast coffee and sat at the kitchen table, spreading the '
        'letter out before her. The handwriting was unmistakably her grandmother\'s, though the woman '
        'had passed away nearly three months ago. "My dearest Elena," it began, "there are things about '
        'our family that I could never tell you while I was alive. The time has come for you to learn '
        'the truth about the Vasquez legacy."'
    )
    doc.add_paragraph(
        'She read the letter three times, each reading revealing new layers of meaning she had initially '
        'missed. Her grandmother spoke of a hidden archive, a collection of documents and artifacts that '
        'had been preserved by the family for over two hundred years. The archive was housed somewhere '
        'in the old estate in Thornfield, a place Elena had not visited since childhood.'
    )
    doc.add_paragraph(
        'By noon, Elena had booked a train ticket to Thornfield. She packed a small suitcase with '
        'essentials: a notebook, a flashlight, her grandmother\'s favorite jade pendant, and a change '
        'of clothes. The journey would take four hours, giving her plenty of time to think about what '
        'she might find when she arrived at the estate.'
    )

    # --- Chapter 2: The Journey ---
    doc.add_page_break()
    h2 = doc.add_heading('Chapter 2: The Journey', level=1)
    doc.add_paragraph(
        'The train from Whitmore Station departed at precisely 2:15 PM, its wheels grinding against '
        'the iron tracks as it pulled away from the platform. Elena found a window seat in the quiet '
        'car and watched the cityscape dissolve into rolling countryside. Fields of wheat stretched '
        'toward the horizon, interrupted only by clusters of oak trees and weathered farmhouses.'
    )
    doc.add_paragraph(
        'She opened her notebook and began writing down everything she remembered about the Thornfield '
        'estate. The main house was a Victorian manor built in 1872 by her great-great-grandfather, '
        'Alejandro Vasquez, who had emigrated from Spain with little more than a trunk of books and '
        'an unwavering determination to build something lasting. The house had fourteen rooms, a '
        'sprawling garden, and a stone cellar that her grandmother had always kept locked.'
    )
    doc.add_paragraph(
        'As the train passed through the town of Millbrook, Elena noticed a fellow passenger watching '
        'her from across the aisle. He was a man in his sixties, with silver hair and sharp, observant '
        'eyes. He wore a tweed jacket and carried an old leather briefcase. When their eyes met, he '
        'offered a slight nod and returned to his newspaper. Elena felt a prickle of unease but '
        'dismissed it as travel nerves.'
    )
    doc.add_paragraph(
        'The conductor announced their approach to Thornfield just as the sun began its descent behind '
        'the western hills. Elena gathered her belongings and stepped onto the platform, breathing in '
        'the familiar scent of pine and damp earth. A taxi was waiting outside the station, its driver '
        'a young woman named Clara who recognized the Vasquez name immediately.'
    )

    # --- Chapter 3: The Discovery ---
    # NO page break here - task says Chapter 3 starts midway on a page
    h3 = doc.add_heading('Chapter 3: The Discovery', level=1)
    doc.add_paragraph(
        'The Thornfield estate loomed before Elena like a sentinel from another era. Its stone facade '
        'was covered in ivy, and the iron gate at the entrance creaked ominously as Clara pushed it '
        'open. The gardens, once meticulously maintained by her grandmother, had grown wild in the '
        'months since her passing. Roses tangled with weeds, and the stone fountain in the center '
        'courtyard stood dry and silent.'
    )
    doc.add_paragraph(
        'Elena used the brass key from the attorney\'s envelope to unlock the front door. The foyer '
        'was exactly as she remembered: a grand staircase ascending to the second floor, portraits of '
        'Vasquez ancestors lining the walls, and the faint smell of lavender that her grandmother had '
        'always kept in small sachets throughout the house. But something was different. On the console '
        'table near the door, someone had placed a fresh bouquet of white lilies.'
    )
    doc.add_paragraph(
        'Following her grandmother\'s coded instructions from the letter, Elena made her way to the '
        'library on the second floor. The room was vast, its walls lined with bookshelves from floor '
        'to ceiling. She located the third shelf from the left, counted seven books from the end, and '
        'pulled out a worn copy of "Don Quixote" in its original Spanish. Behind it, just as the '
        'letter described, was a small brass lever.'
    )
    doc.add_paragraph(
        'The lever released a hidden panel in the wall, revealing a narrow passage that descended into '
        'darkness. Elena switched on her flashlight and stepped through the opening. The passage was '
        'lined with cool stone, and the air grew progressively damper as she descended a spiral '
        'staircase. After what felt like two full rotations, she emerged into a chamber that took '
        'her breath away.'
    )
    doc.add_paragraph(
        'The archive was magnificent. Rows of cedar shelves held leather-bound journals, rolled '
        'parchments, and wooden boxes filled with correspondence dating back to the eighteenth century. '
        'A large oak desk sat in the center of the room, illuminated by a gas lamp that appeared to '
        'still function. On the desk lay a single journal, its cover embossed with the Vasquez family '
        'crest: a silver falcon clutching an olive branch.'
    )

    # --- Chapter 4: The Revelation ---
    doc.add_page_break()
    h4 = doc.add_heading('Chapter 4: The Revelation', level=1)
    doc.add_paragraph(
        'Elena spent the next three days immersed in the archive. She read journals spanning six '
        'generations of the Vasquez family, each one revealing more about the family\'s extraordinary '
        'history. Alejandro Vasquez, it turned out, had not been a simple immigrant. He had been a '
        'cartographer for the Spanish Royal Geographic Society, tasked with mapping unexplored regions '
        'of South America in the 1860s.'
    )
    doc.add_paragraph(
        'During one of his expeditions along the Amazon tributary known as the Rio Escondido, Alejandro '
        'had discovered a series of ancient stone structures hidden beneath centuries of jungle growth. '
        'The structures bore inscriptions in a language that no European scholar could decipher. '
        'Alejandro spent the remaining years of his career quietly documenting the site, convinced '
        'that it held the key to understanding a lost civilization.'
    )
    doc.add_paragraph(
        'The journal on the desk belonged to Elena\'s grandmother, Maria Vasquez, and it contained '
        'her own research into the inscriptions. Maria had been a linguistics professor at the '
        'University of Barcelona before retiring to Thornfield. She had made significant progress '
        'in decoding the symbols, identifying them as a proto-writing system that predated any known '
        'South American civilization by at least two thousand years.'
    )
    doc.add_paragraph(
        'More astonishing still, Maria\'s translations suggested that the structures were not temples '
        'or dwellings, but an elaborate astronomical observatory. The inscriptions described celestial '
        'events with a precision that rivaled modern instruments, including the prediction of eclipses '
        'and the movements of planets invisible to the naked eye.'
    )

    # --- Chapter 5: The Decision ---
    doc.add_page_break()
    h5 = doc.add_heading('Chapter 5: The Decision', level=1)
    doc.add_paragraph(
        'On the fourth morning, Elena sat at the oak desk with Maria\'s journal open before her and '
        'a cup of tea growing cold at her elbow. The weight of what she had learned pressed down on '
        'her like a physical force. Her family had guarded this secret for generations, each member '
        'passing the responsibility to the next. Now the burden fell to her.'
    )
    doc.add_paragraph(
        'She considered her options carefully. She could seal the archive and walk away, preserving '
        'the secret as her ancestors had done. She could contact the academic community and share '
        'Maria\'s research, which would likely spark an international archaeological expedition. Or '
        'she could follow in Alejandro\'s footsteps and travel to the Rio Escondido herself, using '
        'the detailed maps and coordinates preserved in the archive.'
    )
    doc.add_paragraph(
        'That evening, as Elena stood in the garden watching the stars appear one by one above '
        'Thornfield, she made her decision. She would honor both her grandmother\'s scholarship and '
        'her great-great-grandfather\'s spirit of exploration. She would assemble a small team of '
        'trusted researchers and mount a proper expedition to the site. The truth about the Rio '
        'Escondido observatory would be revealed to the world, but carefully and responsibly.'
    )
    doc.add_paragraph(
        'Elena locked the archive, returned the copy of Don Quixote to its place on the shelf, and '
        'closed the hidden panel. She would return to Whitmore, make the necessary arrangements, and '
        'begin planning the journey that would change not only her life but the understanding of human '
        'history itself. As she stepped out into the cool night air, she whispered a quiet thank you '
        'to her grandmother for trusting her with the Vasquez legacy.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
