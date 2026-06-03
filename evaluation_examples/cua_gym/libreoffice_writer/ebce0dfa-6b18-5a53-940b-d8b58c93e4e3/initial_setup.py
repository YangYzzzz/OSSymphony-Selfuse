"""
Initial Setup: team newsletter document for image insertion task
Task ID: writer_obj_065
Domain: libreoffice_writer

Creates:
  - /home/user/Desktop/team_newsletter.docx  (initial state, NO image on page 2)
  - /home/user/Desktop/team_photo.jpg         (2000x1200 pixel placeholder photo)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    from PIL import Image, ImageDraw

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_065'
DOC_OUTPUT = f'{WORKDIR}/team_newsletter.docx'
IMG_OUTPUT = f'{WORKDIR}/team_photo.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_team_photo():
    """Create a realistic 2000x1200 JPEG team photo placeholder."""
    img = Image.new('RGB', (2000, 1200), color=(135, 170, 210))
    draw = ImageDraw.Draw(img)

    # Sky gradient effect
    for y in range(400):
        blue_val = int(180 - y * 0.2)
        draw.line([(0, y), (2000, y)], fill=(100 + y // 10, 150 + y // 20, blue_val))

    # Ground / grass
    for y in range(400, 1200):
        green = int(80 + (y - 400) * 0.05)
        draw.line([(0, y), (2000, y)], fill=(60, min(green, 140), 60))

    # Silhouettes of people (team members)
    people_positions = [
        (200, 700), (350, 680), (500, 710), (650, 690),
        (800, 700), (950, 680), (1100, 710), (1250, 690),
        (1400, 700), (1550, 680), (1700, 710), (1850, 690),
    ]
    for (x, y) in people_positions:
        # Body
        draw.ellipse([x - 25, y - 100, x + 25, y - 50], fill=(50, 40, 35))
        draw.rectangle([x - 20, y - 50, x + 20, y + 60], fill=(70, 100, 150))
        # Legs
        draw.rectangle([x - 20, y + 60, x - 5, y + 120], fill=(40, 40, 40))
        draw.rectangle([x + 5, y + 60, x + 20, y + 120], fill=(40, 40, 40))

    # Banner text area
    draw.rectangle([600, 950, 1400, 1050], fill=(255, 255, 255, 200))
    draw.text((700, 970), "Annual Team Retreat 2025", fill=(30, 30, 30))

    # Decorative border
    draw.rectangle([10, 10, 1990, 1190], outline=(200, 200, 100), width=8)

    img.save(IMG_OUTPUT, 'JPEG', quality=90)
    print(f'Team photo created: {IMG_OUTPUT} (2000x1200)')


def create_newsletter():
    """Create team_newsletter.docx with content on two pages. No image inserted."""
    doc = Document()

    # ── Page 1: Newsletter header and intro ──────────────────────────────────
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Acme Corp Employee Newsletter')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(6)

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run('Volume 8 · Issue 3 · March 2025')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.space_after = Pt(18)

    # Horizontal rule (simulated via border paragraph)
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(0)
    hr_para.paragraph_format.space_after = Pt(12)
    hr_run = hr_para.add_run('─' * 80)
    hr_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Editor's note
    heading1 = doc.add_paragraph()
    h1_run = heading1.add_run("Editor's Note")
    h1_run.bold = True
    h1_run.font.size = Pt(16)
    h1_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    heading1.paragraph_format.space_after = Pt(8)

    doc.add_paragraph(
        "Welcome to the March edition of the Acme Corp newsletter! This has been "
        "an exciting quarter for our team, with major milestones achieved across "
        "all departments. We're proud to share updates on our Q1 performance, "
        "team activities, and upcoming initiatives."
    )

    doc.add_paragraph(
        "Special recognition goes to the Engineering team for successfully launching "
        "the new customer portal on schedule, and to our Sales team for exceeding "
        "targets by 18% in February. These achievements reflect the dedication and "
        "talent of everyone in our organization."
    )

    # Company news section
    heading2 = doc.add_paragraph()
    h2_run = heading2.add_run('Company News & Highlights')
    h2_run.bold = True
    h2_run.font.size = Pt(16)
    h2_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(8)

    doc.add_paragraph(
        "Q1 Financial Results: Our first-quarter results show strong performance "
        "across all business units. Revenue grew by 23% year-over-year, driven "
        "primarily by our enterprise segment and the successful rollout of premium "
        "subscription tiers."
    )

    doc.add_paragraph(
        "New Office Expansion: We are pleased to announce the opening of our new "
        "Singapore office, bringing our total global offices to 12. The Singapore "
        "hub will serve as our Asia-Pacific headquarters, housing teams from Sales, "
        "Customer Success, and Product Engineering."
    )

    doc.add_paragraph(
        "Product Updates: The development team shipped three major feature releases "
        "in Q1. The AI-powered analytics dashboard has been particularly well "
        "received, with over 850 enterprise customers activating it within the "
        "first two weeks of availability."
    )

    # Page break to start page 2
    doc.add_page_break()

    # ── Page 2: Team retreat section ─────────────────────────────────────────
    retreat_heading = doc.add_paragraph()
    rh_run = retreat_heading.add_run('Annual Team Retreat Recap')
    rh_run.bold = True
    rh_run.font.size = Pt(18)
    rh_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    retreat_heading.paragraph_format.space_after = Pt(10)

    # First paragraph on page 2 — this is where the image will be anchored
    first_p2 = doc.add_paragraph(
        "This year's annual team retreat took place over three days at the beautiful "
        "Pines Mountain Resort in Lake Tahoe. More than 120 employees gathered from "
        "offices across North America, Europe, and Asia to connect, collaborate, and "
        "celebrate our collective achievements. The retreat theme, 'Growing Together,' "
        "set the tone for a weekend of meaningful conversations and shared experiences."
    )
    first_p2.paragraph_format.space_after = Pt(10)

    doc.add_paragraph(
        "Day one kicked off with a keynote from CEO Rachel Nguyen, who reflected on the "
        "company's journey from a 12-person startup to a 600-strong global organization. "
        "Her talk highlighted the values that have guided our growth: transparency, "
        "curiosity, and a relentless focus on customer impact. The session was followed "
        "by department showcases, where teams demonstrated their Q1 wins and shared "
        "upcoming roadmaps."
    )

    doc.add_paragraph(
        "Workshop highlights included a design-thinking sprint facilitated by the "
        "Product team, resulting in three promising concepts for new features. The "
        "Leadership Development track drew 40 participants for sessions on coaching, "
        "feedback culture, and building psychological safety in remote-first teams. "
        "Several managers noted that the frameworks shared were immediately applicable "
        "to their day-to-day work."
    )

    doc.add_paragraph(
        "Evening activities brought the fun — from a lakeside barbecue on Friday night "
        "to a team talent show on Saturday that revealed surprising hidden talents (we "
        "had no idea our CFO, David Park, was an accomplished folk guitarist!). The "
        "annual 5K fun run on Sunday morning drew 65 participants, with Marketing's "
        "Elena Kovacs taking the top spot for the second year running."
    )

    doc.add_paragraph(
        "Feedback from attendees was overwhelmingly positive, with 94% rating the event "
        "as 'excellent' or 'outstanding.' Planning for next year's retreat has already "
        "begun, with a committee of volunteers from across the company working on venue "
        "options and programming ideas. Watch for updates in the coming months!"
    )

    # Save the document
    doc.save(DOC_OUTPUT)
    print(f'Newsletter created: {DOC_OUTPUT}')


def create_initial():
    """Run all setup steps."""
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    # 1. Create the team photo JPEG
    create_team_photo()

    # 2. Create the newsletter document (no image)
    create_newsletter()

    # 3. GUI startup — open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
