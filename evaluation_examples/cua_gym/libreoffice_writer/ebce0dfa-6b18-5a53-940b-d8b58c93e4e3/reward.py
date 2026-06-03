"""
Reward Script: Insert and configure team_photo.jpg image in team_newsletter.docx
Task ID: writer_obj_065
Domain: libreoffice_writer
Scoring:
  Component 1: Image is present in the document (0.20 pts)
  Component 2: Image visible width is 10cm / cx=3600000 EMU (0.20 pts)
  Component 3: Crop values match task spec (top=2cm, bottom=1cm, left=0, right=0) (0.25 pts)
  Component 4: Wrapping is Parallel (wrapSquare/bothSides) with 0.4cm spacing all sides (0.20 pts)
  Component 5: Alt text is 'Team photo from annual retreat 2025' (0.10 pts)
  Component 6: Image anchored to first paragraph on page 2 (0.05 pts)
  Total: 1.00 pts
"""

import os

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_065'

# EMU conversion: 1cm = 360000 EMU
CM_TO_EMU = 360000

# Tolerances
EMU_TOLERANCE = 18000    # 0.05cm tolerance on dimensions
CROP_TOLERANCE = 200     # tolerance in 1/100000 units for srcRect values

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document
    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find drawing elements (anchors/inline) in the document body
    try:
        import lxml.etree as etree
        ns = {
            'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
        body = doc.element.body
        anchor_drawings = body.findall('.//wp:anchor', ns)
        inline_drawings = body.findall('.//wp:inline', ns)
        all_drawings = anchor_drawings + inline_drawings
    except Exception as e:
        print(f"ERROR: Cannot parse document XML: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Image is present in the document (0.20 pts)
    try:
        img_rels = [rel for rel in doc.part.rels.values() if "image" in rel.reltype]
        has_image = len(img_rels) > 0 and len(all_drawings) > 0
        if has_image:
            print(f"PASS: Component 1 — Image found in document (rels={len(img_rels)}, drawings={len(all_drawings)}) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — No image found in document (rels={len(img_rels)}, drawings={len(all_drawings)})")
            # No image = task not done; return early
            print(f"\nScore: {total_score}/1.0")
            print(f"REWARD: {total_score}")
            return total_score
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Work with the first drawing element
    drawing = all_drawings[0]

    # Component 2: Image visible width is 10cm (cx = 3600000 EMU) (0.20 pts)
    try:
        extent = drawing.find('.//wp:extent', ns)
        if extent is None:
            # Try pic:spPr extent
            ext_elem = drawing.find('.//a:ext', ns)
        else:
            ext_elem = extent

        if ext_elem is not None:
            cx = int(ext_elem.get('cx', 0))
            cy = int(ext_elem.get('cy', 0))
        else:
            cx = 0
            cy = 0

        expected_cx = 10 * CM_TO_EMU  # 3600000 EMU = 10cm
        if abs(cx - expected_cx) <= EMU_TOLERANCE:
            print(f"PASS: Component 2 — Image width is 10cm (cx={cx}, expected={expected_cx}, tolerance={EMU_TOLERANCE}) (0.20 pts)")
            total_score += 0.20
        else:
            actual_cm = cx / CM_TO_EMU
            print(f"FAIL: Component 2 — Expected width 10cm (cx={expected_cx}), found {actual_cm:.3f}cm (cx={cx})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Crop values match task spec (top=2cm, bottom=1cm, left=0, right=0) (0.25 pts)
    # srcRect values are in 1/100000 of total image dimension
    # Image is 2000x1200px; at standard resolution (96 DPI), natural height in cm depends on DPI
    # The actual crop fractions stored in srcRect are derived from: top_fraction = crop_cm / natural_height_cm
    # From golden analysis: t=6299 (6.299% = 2cm crop), b=3150 (3.15% = 1cm crop)
    # Expected t ≈ 6299, b ≈ 3150, l=0, r=0
    EXPECTED_T = 6299
    EXPECTED_B = 3150
    EXPECTED_L = 0
    EXPECTED_R = 0
    try:
        src_rect = drawing.find('.//a:srcRect', ns)
        if src_rect is not None:
            t_val = int(src_rect.get('t', '0'))
            b_val = int(src_rect.get('b', '0'))
            l_val = int(src_rect.get('l', '0'))
            r_val = int(src_rect.get('r', '0'))

            top_ok = abs(t_val - EXPECTED_T) <= CROP_TOLERANCE
            bot_ok = abs(b_val - EXPECTED_B) <= CROP_TOLERANCE
            left_ok = abs(l_val - EXPECTED_L) <= CROP_TOLERANCE
            right_ok = abs(r_val - EXPECTED_R) <= CROP_TOLERANCE

            if top_ok and bot_ok and left_ok and right_ok:
                print(f"PASS: Component 3 — Crop values correct: t={t_val}, b={b_val}, l={l_val}, r={r_val} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — Crop mismatch.")
                if not top_ok:
                    print(f"  Top: expected ~{EXPECTED_T} (2cm), found {t_val}")
                if not bot_ok:
                    print(f"  Bottom: expected ~{EXPECTED_B} (1cm), found {b_val}")
                if not left_ok:
                    print(f"  Left: expected {EXPECTED_L}, found {l_val}")
                if not right_ok:
                    print(f"  Right: expected {EXPECTED_R}, found {r_val}")
        else:
            print(f"FAIL: Component 3 — No srcRect element found (crop not applied)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Wrapping is Parallel (wrapSquare/bothSides) with 0.4cm spacing all sides (0.20 pts)
    # 0.4cm = 0.4 * 360000 = 144000 EMU
    EXPECTED_DIST = 144000  # 0.4cm in EMU
    DIST_TOLERANCE = 7200   # 0.02cm tolerance
    try:
        # Check anchor element for dist attributes (both-sides wrapping sets these on anchor itself)
        wrap_square = drawing.find('.//wp:wrapSquare', ns)
        wrap_tight = drawing.find('.//wp:wrapTight', ns)
        wrap_through = drawing.find('.//wp:wrapThrough', ns)

        # "Parallel" wrapping in LibreOffice typically maps to wrapSquare with bothSides
        has_parallel_wrap = (wrap_square is not None and
                             wrap_square.get('wrapText', '') in ('bothSides', 'largest'))

        if not has_parallel_wrap:
            # Also accept wrapSquare without wrapText attribute check (default is bothSides)
            has_parallel_wrap = wrap_square is not None

        if has_parallel_wrap:
            # Check spacing (dist attributes on anchor element or wrapSquare)
            anchor_elem = drawing  # The drawing IS the anchor
            distT = int(anchor_elem.get('distT', 0))
            distB = int(anchor_elem.get('distB', 0))
            distL = int(anchor_elem.get('distL', 0))
            distR = int(anchor_elem.get('distR', 0))

            # Also check on wrapSquare itself
            if distT == 0 and distB == 0 and distL == 0 and distR == 0:
                distT = int(wrap_square.get('distT', 0))
                distB = int(wrap_square.get('distB', 0))
                distL = int(wrap_square.get('distL', 0))
                distR = int(wrap_square.get('distR', 0))

            dist_ok = (abs(distT - EXPECTED_DIST) <= DIST_TOLERANCE and
                       abs(distB - EXPECTED_DIST) <= DIST_TOLERANCE and
                       abs(distL - EXPECTED_DIST) <= DIST_TOLERANCE and
                       abs(distR - EXPECTED_DIST) <= DIST_TOLERANCE)

            if dist_ok:
                print(f"PASS: Component 4 — Parallel wrapping with 0.4cm spacing: "
                      f"distT={distT}, distB={distB}, distL={distL}, distR={distR} (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — Wrapping is Parallel but spacing is wrong. "
                      f"Expected {EXPECTED_DIST} EMU (0.4cm) on all sides. "
                      f"Found: distT={distT}, distB={distB}, distL={distL}, distR={distR}")
        else:
            wrap_type = "none"
            if drawing.find('.//wp:wrapNone', ns) is not None:
                wrap_type = "wrapNone"
            elif drawing.find('.//wp:wrapTopAndBottom', ns) is not None:
                wrap_type = "wrapTopAndBottom"
            print(f"FAIL: Component 4 — Expected Parallel wrapping (wrapSquare), found: {wrap_type}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Alt text is 'Team photo from annual retreat 2025' (0.10 pts)
    EXPECTED_ALT_TEXT = "Team photo from annual retreat 2025"
    try:
        doc_pr = drawing.find('.//wp:docPr', ns)
        if doc_pr is not None:
            descr = doc_pr.get('descr', '')
            if descr.strip() == EXPECTED_ALT_TEXT:
                print(f"PASS: Component 5 — Alt text correct: {descr!r} (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 5 — Alt text mismatch. Expected: {EXPECTED_ALT_TEXT!r}, Found: {descr!r}")
        else:
            print(f"FAIL: Component 5 — No docPr element found (alt text not set)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Image is anchored to first paragraph on page 2 (0.05 pts)
    # Page 2 starts after a page break. The first paragraph on page 2 is "Annual Team Retreat Recap".
    # The anchor drawing should be in that paragraph.
    try:
        from docx.oxml.ns import qn

        # Find the paragraph containing a page break to determine page boundary
        page_break_para_idx = None
        for i, para in enumerate(doc.paragraphs):
            para_elem = para._element
            brs = para_elem.findall('.//w:br', ns)
            for br in brs:
                br_type = br.get(qn('w:type'))
                if br_type == 'page':
                    page_break_para_idx = i
                    break
            if page_break_para_idx is not None:
                break

        if page_break_para_idx is None:
            print("FAIL: Component 6 — No page break found; cannot determine page 2 start")
        else:
            # The first paragraph on page 2 is the one after the page break paragraph
            first_p2_idx = page_break_para_idx + 1
            paras = doc.paragraphs
            if first_p2_idx < len(paras):
                first_p2_para = paras[first_p2_idx]
                # Check if this paragraph contains the anchor
                para_anchors = first_p2_para._element.findall('.//wp:anchor', ns)
                para_inlines = first_p2_para._element.findall('.//wp:inline', ns)
                if para_anchors or para_inlines:
                    print(f"PASS: Component 6 — Image anchored to first paragraph on page 2 "
                          f"(para {first_p2_idx}: {first_p2_para.text[:60]!r}) (0.05 pts)")
                    total_score += 0.05
                else:
                    print(f"FAIL: Component 6 — Image NOT in first paragraph on page 2 "
                          f"(para {first_p2_idx}: {first_p2_para.text[:60]!r}). "
                          f"Check which paragraph contains the drawing.")
            else:
                print(f"FAIL: Component 6 — page_break_para_idx={page_break_para_idx}, first_p2_idx={first_p2_idx} out of range")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/team_newsletter.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
