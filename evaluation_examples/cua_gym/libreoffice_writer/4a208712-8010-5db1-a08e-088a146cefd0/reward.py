"""
Reward Script: Format job description with consistent styling
Task ID: writer_hr_028
Domain: libreoffice_writer
Scoring:
  Component 1 (0.2): Job title 'Product Manager' is Heading 1
  Component 2 (0.3): Section labels are Heading 2
  Component 3 (0.3): Bullet items use List Bullet style
  Component 4 (0.2): Bullet items have ~0.5 inch left indent
"""

import os
from docx import Document
from docx.shared import Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_028'

# 0.5 inch in EMU = 457200
EXPECTED_INDENT_EMU = 457200
INDENT_TOLERANCE = 50000  # ~0.04 inch tolerance

# Section labels that should be Heading 2
SECTION_LABELS = {'Responsibilities', 'Qualifications', 'Benefits', 'About Us'}


def persist_app_state():
    """Try to save any unsaved LibreOffice Writer state."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def classify_paragraphs(doc):
    """Classify paragraphs into title, section labels, bullet items, and blanks."""
    title_idx = None
    section_indices = []
    bullet_indices = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if text == 'Product Manager' and title_idx is None:
            title_idx = i
        elif text in SECTION_LABELS:
            section_indices.append(i)
        else:
            # Non-empty, non-title, non-section-label => should be a bullet item
            bullet_indices.append(i)

    return title_idx, section_indices, bullet_indices


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    title_idx, section_indices, bullet_indices = classify_paragraphs(doc)

    # Component 1: Job title 'Product Manager' is Heading 1 (0.2 points)
    try:
        if title_idx is not None:
            style_name = doc.paragraphs[title_idx].style.name
            if style_name == 'Heading 1':
                print(f"PASS: Component 1 — 'Product Manager' has style 'Heading 1' (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 1 — 'Product Manager' style is '{style_name}', expected 'Heading 1'")
        else:
            print("FAIL: Component 1 — 'Product Manager' paragraph not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 4 section labels are Heading 2 (0.3 points)
    try:
        if len(section_indices) == 0:
            print("FAIL: Component 2 — No section label paragraphs found")
        else:
            heading2_count = 0
            for idx in section_indices:
                style_name = doc.paragraphs[idx].style.name
                label = doc.paragraphs[idx].text.strip()
                if style_name == 'Heading 2':
                    heading2_count += 1
                else:
                    print(f"  DETAIL: '{label}' has style '{style_name}', expected 'Heading 2'")

            ratio = heading2_count / len(section_indices)
            points = round(0.3 * ratio, 2)
            if heading2_count == len(section_indices):
                print(f"PASS: Component 2 — All {heading2_count} section labels are Heading 2 ({points} pts)")
                total_score += points
            elif heading2_count > 0:
                print(f"PARTIAL: Component 2 — {heading2_count}/{len(section_indices)} section labels are Heading 2 ({points} pts)")
                total_score += points
            else:
                print(f"FAIL: Component 2 — 0/{len(section_indices)} section labels are Heading 2")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Bullet items use 'List Bullet' style (0.3 points)
    try:
        if len(bullet_indices) == 0:
            print("FAIL: Component 3 — No bullet item paragraphs found")
        else:
            list_bullet_count = 0
            for idx in bullet_indices:
                style_name = doc.paragraphs[idx].style.name
                if style_name == 'List Bullet':
                    list_bullet_count += 1

            ratio = list_bullet_count / len(bullet_indices)
            points = round(0.3 * ratio, 2)
            if list_bullet_count == len(bullet_indices):
                print(f"PASS: Component 3 — All {list_bullet_count} bullet items use 'List Bullet' style ({points} pts)")
                total_score += points
            elif list_bullet_count > 0:
                print(f"PARTIAL: Component 3 — {list_bullet_count}/{len(bullet_indices)} bullet items use 'List Bullet' style ({points} pts)")
                total_score += points
            else:
                print(f"FAIL: Component 3 — 0/{len(bullet_indices)} bullet items use 'List Bullet' style")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Bullet items have ~0.5 inch (457200 EMU) left indent (0.2 points)
    try:
        if len(bullet_indices) == 0:
            print("FAIL: Component 4 — No bullet item paragraphs found")
        else:
            indent_ok_count = 0
            for idx in bullet_indices:
                indent = doc.paragraphs[idx].paragraph_format.left_indent
                if indent is not None and abs(indent - EXPECTED_INDENT_EMU) <= INDENT_TOLERANCE:
                    indent_ok_count += 1

            ratio = indent_ok_count / len(bullet_indices)
            points = round(0.2 * ratio, 2)
            if indent_ok_count == len(bullet_indices):
                print(f"PASS: Component 4 — All {indent_ok_count} bullet items have correct left indent ({points} pts)")
                total_score += points
            elif indent_ok_count > 0:
                print(f"PARTIAL: Component 4 — {indent_ok_count}/{len(bullet_indices)} bullet items have correct left indent ({points} pts)")
                total_score += points
            else:
                print(f"FAIL: Component 4 — 0/{len(bullet_indices)} bullet items have correct left indent")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 1)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
