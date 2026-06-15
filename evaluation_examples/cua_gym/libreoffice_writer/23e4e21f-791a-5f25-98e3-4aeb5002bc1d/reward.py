"""
Reward Script: Accept all tracked changes from 'Partner_Review' and reject all from 'Associate_Draft'
Task ID: writer_legal_055
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): No tracked changes remain in the document
  Component 2 (0.35): Partner_Review changes accepted — specific paragraph text matches golden
  Component 3 (0.20): Associate_Draft insertions rejected — inserted text NOT present
  Component 4 (0.25): Associate_Draft deletions rejected — original text restored
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_055'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for %s" % domain)
        except Exception as e:
            print("PERSIST_WARN: save hook failed: %s" % e)


def get_full_para_text_with_revisions(para_element, ns):
    """
    Get the full text of a paragraph including accepted tracked changes.
    This returns text as it would appear if all insertions are accepted and
    deletions removed — i.e., what para.text gives. We need this to understand
    that para.text already incorporates w:ins text but excludes w:del text.
    """
    # para.text in python-docx already does the right thing for non-tracked docs
    # For tracked-change docs, we need to read XML directly
    pass


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Build paragraph text list
    para_texts = [p.text for p in doc.paragraphs]

    # =========================================================================
    # Component 1: No tracked changes remain (0.20 points)
    # After accepting Partner_Review and rejecting Associate_Draft, zero
    # revision marks should remain in the XML.
    # This check FAILS on initial (20 changes present) and PASSES on golden (0).
    # =========================================================================
    try:
        body = doc.element.body
        insertions = body.findall('.//w:ins', ns)
        deletions = body.findall('.//w:del', ns)
        total_changes = len(insertions) + len(deletions)
        if total_changes == 0:
            print("PASS: Component 1 — No tracked changes remain (0.20 pts)")
            total_score += 0.20
        else:
            print("FAIL: Component 1 — Found %d tracked changes remaining (ins=%d, del=%d)" %
                  (total_changes, len(insertions), len(deletions)))
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # =========================================================================
    # Component 2: Partner_Review changes accepted — paragraph text matches (0.35 pts)
    # After accepting Partner_Review edits, specific paragraphs should contain
    # the inserted text and NOT contain the deleted text.
    # We check PARAGRAPH-SPECIFIC text to avoid false positives from other paras.
    #
    # In initial file: para.text EXCLUDES both w:ins and w:del content from
    # tracked changes, so these checks fail on initial.
    # In golden file: tracked changes resolved, para.text shows final content.
    # =========================================================================
    try:
        pr_checks = [
            # (para_idx, must_contain, must_not_contain, description)
            (2, "duly organized and", None,
             "P2: 'duly organized and' present in corporation description"),
            (5, "comprehensive", "basic",
             "P5: 'comprehensive' present, 'basic' absent"),
            (9, "$200", "$150",
             "P9: '$200' rate present, '$150' absent"),
            (11, "thirty (30)", "sixty (60)",
             "P11: 'thirty (30)' present, 'sixty (60)' absent"),
            (20, "any third party without prior written consent", None,
             "P20: full confidentiality clause present"),
            (24, "work made for hire", None,
             "P24: 'work made for hire' IP clause present"),
            (34, "Delaware", "California",
             "P34: 'Delaware' present, 'California' absent"),
        ]
        passed = 0
        for para_idx, must_have, must_not_have, desc in pr_checks:
            if para_idx >= len(para_texts):
                print("  WARN: Component 2 — para index %d out of range" % para_idx)
                continue
            ptext = para_texts[para_idx]
            has_required = must_have in ptext if must_have else True
            lacks_forbidden = (must_not_have not in ptext) if must_not_have else True
            if has_required and lacks_forbidden:
                passed += 1
                print("  PASS: Component 2 — %s" % desc)
            else:
                reasons = []
                if not has_required:
                    reasons.append("'%s' not found" % must_have)
                if not lacks_forbidden:
                    reasons.append("'%s' still present" % must_not_have)
                print("  FAIL: Component 2 — %s — %s" % (desc, ', '.join(reasons)))

        if passed == len(pr_checks):
            print("PASS: Component 2 — All %d Partner_Review changes accepted (0.35 pts)" % passed)
            total_score += 0.35
        elif passed > 0:
            partial = round(0.35 * passed / len(pr_checks), 3)
            print("PARTIAL: Component 2 — %d/%d checks passed (%.3f pts)" % (passed, len(pr_checks), partial))
            total_score += partial
        else:
            print("FAIL: Component 2 — No Partner_Review changes detected")
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # =========================================================================
    # Component 3: Associate_Draft insertions rejected — text NOT present (0.20 pts)
    # These texts were inserted by Associate_Draft. Rejecting means they should
    # NOT be in the final document.
    # PARAGRAPH-SPECIFIC checks to avoid false matches.
    #
    # In initial file: para.text EXCLUDES tracked insertion text, so these
    # "must not contain" checks would pass on initial too.
    # FIX: We combine with a POSITIVE check — the paragraph must contain its
    # expected golden content (which differs from initial blank text).
    # For P37 (Notwithstanding clause): we check the raw XML for w:ins elements
    # since para.text is identical in both initial and golden for that paragraph.
    # =========================================================================
    try:
        ad_ins_checks = [
            # (para_idx, must_not_contain, must_contain_instead, description)
            (2, "formerly known as Apex Solutions",
             "duly organized and",
             "P2: 'formerly known as Apex Solutions' absent AND 'duly organized and' present"),
            (14, "two (2) years",
             "one (1) year",
             "P14: 'two (2) years' absent AND 'one (1) year' present"),
            (21, "two (2) years",
             "five (5) years",
             "P21: 'two (2) years' absent AND 'five (5) years' present"),
        ]
        passed = 0
        for para_idx, must_not_have, must_have, desc in ad_ins_checks:
            if para_idx >= len(para_texts):
                print("  WARN: Component 3 — para index %d out of range" % para_idx)
                continue
            ptext = para_texts[para_idx]
            lacks_forbidden = must_not_have not in ptext
            has_required = must_have in ptext if must_have else True
            if lacks_forbidden and has_required:
                passed += 1
                print("  PASS: Component 3 — %s" % desc)
            else:
                reasons = []
                if not lacks_forbidden:
                    reasons.append("'%s' found (should be absent)" % must_not_have)
                if not has_required:
                    reasons.append("'%s' not found (should be present)" % must_have)
                print("  FAIL: Component 3 — %s — %s" % (desc, ', '.join(reasons)))

        # Special check for P37: verify no w:ins elements with "Notwithstanding" in XML
        # This catches the tracked insertion in initial but not in golden (where it's removed)
        if len(doc.paragraphs) > 37:
            p37_elem = doc.paragraphs[37]._element
            p37_ins_elements = p37_elem.findall('.//w:ins', ns)
            p37_ins_texts = []
            for ins_el in p37_ins_elements:
                for t in ins_el.findall('.//w:t', ns):
                    if t.text:
                        p37_ins_texts.append(t.text)
            p37_has_notwithstanding_ins = any("Notwithstanding" in t for t in p37_ins_texts)
            # Also check para.text for the inserted text (present if change was accepted, which is wrong)
            p37_has_notwithstanding_text = "Notwithstanding" in para_texts[37]
            if not p37_has_notwithstanding_ins and not p37_has_notwithstanding_text:
                passed += 1
                print("  PASS: Component 3 — P37: no 'Notwithstanding' insertion in XML or text")
            else:
                if p37_has_notwithstanding_ins:
                    print("  FAIL: Component 3 — P37: tracked insertion with 'Notwithstanding' still in XML")
                if p37_has_notwithstanding_text:
                    print("  FAIL: Component 3 — P37: 'Notwithstanding' text accepted into paragraph")

        total_checks = 4  # 3 text checks + 1 XML check
        if passed == total_checks:
            print("PASS: Component 3 — All %d Associate_Draft insertions rejected (0.20 pts)" % passed)
            total_score += 0.20
        elif passed > 0:
            partial = round(0.20 * passed / total_checks, 3)
            print("PARTIAL: Component 3 — %d/%d checks passed (%.3f pts)" % (passed, total_checks, partial))
            total_score += partial
        else:
            print("FAIL: Component 3 — No Associate_Draft insertions rejected")
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    # =========================================================================
    # Component 4: Associate_Draft deletions rejected — original text restored (0.25 pts)
    # These texts were deleted by Associate_Draft. Rejecting means the original
    # text should remain.
    #
    # In initial file: para.text EXCLUDES tracked deletion text (w:delText),
    # so "at Consultant's discretion", "one (1) year", etc. are absent.
    # In golden file: tracked changes resolved, rejected deletions restore text.
    # These checks FAIL on initial and PASS on golden. Good.
    # =========================================================================
    try:
        ad_del_checks = [
            # (para_idx, must_contain, description)
            (6, "at Consultant's discretion",
             "P6: 'at Consultant's discretion' retained in services location"),
            (14, "one (1) year",
             "P14: 'one (1) year' term retained"),
            (21, "five (5) years",
             "P21: 'five (5) years' confidentiality period retained"),
            (26, "irrevocable",
             "P26: 'irrevocable' in IP license retained"),
        ]
        passed = 0
        for para_idx, must_have, desc in ad_del_checks:
            if para_idx >= len(para_texts):
                print("  WARN: Component 4 — para index %d out of range" % para_idx)
                continue
            ptext = para_texts[para_idx]
            if must_have in ptext:
                passed += 1
                print("  PASS: Component 4 — %s" % desc)
            else:
                print("  FAIL: Component 4 — %s — '%s' not found in para" % (desc, must_have))

        if passed == len(ad_del_checks):
            print("PASS: Component 4 — All %d Associate_Draft deletions rejected (0.25 pts)" % passed)
            total_score += 0.25
        elif passed > 0:
            partial = round(0.25 * passed / len(ad_del_checks), 3)
            print("PARTIAL: Component 4 — %d/%d checks passed (%.3f pts)" % (passed, len(ad_del_checks), partial))
            total_score += partial
        else:
            print("FAIL: Component 4 — No Associate_Draft deletions rejected")
    except Exception as e:
        print("ERROR: Component 4 — %s" % e)

    final_score = round(min(total_score, 1.0), 2)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = '%s/%s.docx' % (WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
