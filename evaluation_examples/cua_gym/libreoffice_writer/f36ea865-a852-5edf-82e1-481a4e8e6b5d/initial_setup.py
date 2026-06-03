"""
Initial Setup: Employee handbook without cover page
Task ID: writer_hr_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_051'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ---- Table of Contents ----
    heading_toc = doc.add_heading('Table of Contents', level=1)

    toc_items = [
        '1. Welcome Message',
        '2. Company Overview',
        '3. Employment Policies',
        '4. Code of Conduct',
        '5. Compensation and Benefits',
        '6. Leave Policies',
        '7. Health and Safety',
        '8. Disciplinary Procedures',
        '9. Acknowledgment',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ---- 1. Welcome Message ----
    doc.add_heading('1. Welcome Message', level=1)
    doc.add_paragraph(
        'Welcome to Summit Financial Group! We are thrilled to have you join our team. '
        'This handbook has been designed to help you navigate your employment journey with us. '
        'Whether you are a new hire or a returning team member, we encourage you to familiarize '
        'yourself with the policies and expectations outlined in the following pages.'
    )
    doc.add_paragraph(
        'Our mission is to deliver exceptional financial services while fostering a workplace '
        'culture rooted in integrity, collaboration, and continuous improvement. We believe '
        'that our employees are the foundation of our success, and we are committed to '
        'providing a supportive environment where everyone can thrive.'
    )

    doc.add_page_break()

    # ---- 2. Company Overview ----
    doc.add_heading('2. Company Overview', level=1)
    doc.add_paragraph(
        'Summit Financial Group was founded in 1998 and has grown to become one of the '
        'leading regional financial services firms in the Mid-Atlantic region. Headquartered '
        'in Arlington, Virginia, we employ over 1,200 professionals across 15 offices.'
    )
    doc.add_paragraph(
        'Our core divisions include Wealth Management, Commercial Banking, Insurance Services, '
        'and Financial Advisory. Each division operates under a unified commitment to client '
        'satisfaction and regulatory compliance.'
    )

    doc.add_page_break()

    # ---- 3. Employment Policies ----
    doc.add_heading('3. Employment Policies', level=1)

    doc.add_heading('3.1 Equal Employment Opportunity', level=2)
    doc.add_paragraph(
        'Summit Financial Group is an equal opportunity employer. We do not discriminate '
        'on the basis of race, color, religion, sex, national origin, age, disability, '
        'genetic information, sexual orientation, gender identity, or any other characteristic '
        'protected by federal, state, or local law.'
    )

    doc.add_heading('3.2 At-Will Employment', level=2)
    doc.add_paragraph(
        'Employment with Summit Financial Group is on an at-will basis. This means that '
        'either the employee or the company may terminate the employment relationship at '
        'any time, with or without cause or notice, subject to applicable law.'
    )

    doc.add_heading('3.3 Work Hours and Attendance', level=2)
    doc.add_paragraph(
        'Standard business hours are Monday through Friday, 8:30 AM to 5:00 PM. Employees '
        'are expected to maintain regular attendance and report any absences to their direct '
        'supervisor as soon as practicable. Flexible work arrangements may be available upon '
        'approval by department management.'
    )

    doc.add_page_break()

    # ---- 4. Code of Conduct ----
    doc.add_heading('4. Code of Conduct', level=1)
    doc.add_paragraph(
        'All employees are expected to conduct themselves in a professional manner that '
        'reflects positively on Summit Financial Group. Key expectations include:'
    )
    doc.add_paragraph('Maintaining confidentiality of client and proprietary information', style='List Bullet')
    doc.add_paragraph('Complying with all applicable laws, regulations, and internal policies', style='List Bullet')
    doc.add_paragraph('Avoiding conflicts of interest and reporting potential conflicts promptly', style='List Bullet')
    doc.add_paragraph('Treating colleagues, clients, and partners with respect and dignity', style='List Bullet')
    doc.add_paragraph('Using company resources responsibly and for authorized purposes only', style='List Bullet')

    doc.add_page_break()

    # ---- 5. Compensation and Benefits ----
    doc.add_heading('5. Compensation and Benefits', level=1)

    doc.add_heading('5.1 Pay Schedule', level=2)
    doc.add_paragraph(
        'Employees are paid on a bi-weekly basis, with 26 pay periods per year. Direct '
        'deposit is available and encouraged. Pay stubs can be accessed through the employee '
        'self-service portal at hr.summitfinancial.com.'
    )

    doc.add_heading('5.2 Health Insurance', level=2)
    doc.add_paragraph(
        'Full-time employees are eligible for medical, dental, and vision coverage beginning '
        'the first day of the month following 30 days of employment. Summit Financial Group '
        'covers 80% of employee premiums and 50% of dependent premiums.'
    )

    doc.add_heading('5.3 Retirement Plan', level=2)
    doc.add_paragraph(
        'Employees may participate in the company 401(k) plan after 90 days of employment. '
        'The company provides a dollar-for-dollar match on the first 4% of eligible compensation '
        'contributed by the employee. Vesting is on a three-year graded schedule.'
    )

    doc.add_page_break()

    # ---- 6. Leave Policies ----
    doc.add_heading('6. Leave Policies', level=1)

    doc.add_heading('6.1 Paid Time Off (PTO)', level=2)
    doc.add_paragraph(
        'Full-time employees accrue PTO based on years of service: 0-2 years: 15 days; '
        '3-5 years: 20 days; 6+ years: 25 days. PTO must be requested through the HR '
        'portal and approved by the employee\'s supervisor at least two weeks in advance '
        'for planned absences.'
    )

    doc.add_heading('6.2 Sick Leave', level=2)
    doc.add_paragraph(
        'Employees receive 10 paid sick days per calendar year. Unused sick days may be '
        'carried over up to a maximum balance of 30 days. A physician\'s note may be '
        'required for absences exceeding three consecutive days.'
    )

    doc.add_page_break()

    # ---- 7. Health and Safety ----
    doc.add_heading('7. Health and Safety', level=1)
    doc.add_paragraph(
        'Summit Financial Group is committed to maintaining a safe and healthy workplace. '
        'All employees are responsible for following safety procedures and reporting hazards '
        'or incidents to the Facilities Management team immediately. Emergency evacuation '
        'plans are posted on every floor and reviewed during annual safety drills.'
    )

    doc.add_page_break()

    # ---- 8. Disciplinary Procedures ----
    doc.add_heading('8. Disciplinary Procedures', level=1)
    doc.add_paragraph(
        'The company follows a progressive discipline process for policy violations:'
    )
    doc.add_paragraph('Step 1: Verbal Warning', style='List Number')
    doc.add_paragraph('Step 2: Written Warning', style='List Number')
    doc.add_paragraph('Step 3: Final Written Warning / Suspension', style='List Number')
    doc.add_paragraph('Step 4: Termination', style='List Number')
    doc.add_paragraph(
        'Severe violations, including but not limited to fraud, harassment, theft, or '
        'violence, may result in immediate termination without prior warnings.'
    )

    doc.add_page_break()

    # ---- 9. Acknowledgment ----
    doc.add_heading('9. Acknowledgment', level=1)
    doc.add_paragraph(
        'By signing below, I acknowledge that I have received, read, and understand the '
        'Summit Financial Group Employee Handbook. I agree to comply with the policies and '
        'procedures described herein.'
    )

    # Signature lines
    doc.add_paragraph('')
    doc.add_paragraph('_______________________________________')
    p = doc.add_paragraph('Employee Signature')
    p.paragraph_format.space_after = Pt(24)

    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Date')

    doc.add_paragraph('')
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('HR Representative Signature')
    p2 = doc.add_paragraph('')
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Date')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
