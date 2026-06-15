"""
Reward Script: Apply Heading styles and first-line indents to policy_document.docx
Task ID: writer_para_055
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.4 pts): Paragraph 1 has style 'Heading 1'
  - Component 2 (0.3 pts): Paragraphs 3 and 6 have style 'Heading 2'
  - Component 3 (0.3 pts): Body paragraphs (2,4,5,7,8) have first-line indent ~0.75 cm
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_para_055'

# 0.75 cm in EMU = 270000 (Cm(0.75)); allow ±2000 EMU tolerance for rounding
TARGET_INDENT_EMU = 270000
INDENT_TOLERANCE_EMU = 2000


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have exactly 8 paragraphs with expected text content
    try:
        paras = doc.paragraphs
        if len(paras) != 8:
            print(f"FAIL: Expected 8 paragraphs, found {len(paras)}")
            print("REWARD: 0.0")
            return 0.0
        # Verify key text to confirm correct file
        expected_texts = {
            1: 'Information Security Policy',
            3: 'Purpose and Scope',
            6: 'Password Requirements',
        }
        for idx, expected in expected_texts.items():
            actual = paras[idx - 1].text.strip()
            if actual != expected:
                print(f"FAIL: Paragraph {idx} text mismatch. Expected '{expected}', got '{actual}'")
                print("REWARD: 0.0")
                return 0.0
        print("PRECONDITION: File structure and text content verified.")
    except Exception as e:
        print(f"CRITICAL: Precondition check failed: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Paragraph 1 has style 'Heading 1' (0.4 points)
    try:
        para1 = doc.paragraphs[0]
        style1 = para1.style.name if para1.style else 'None'
        if style1 == 'Heading 1':
            print(f"PASS: Component 1 — Paragraph 1 style is 'Heading 1' (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Paragraph 1 expected style 'Heading 1', found '{style1}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraphs 3 and 6 both have style 'Heading 2' (0.3 points)
    try:
        para3 = doc.paragraphs[2]
        para6 = doc.paragraphs[5]
        style3 = para3.style.name if para3.style else 'None'
        style6 = para6.style.name if para6.style else 'None'
        p3_ok = (style3 == 'Heading 2')
        p6_ok = (style6 == 'Heading 2')
        if p3_ok and p6_ok:
            print(f"PASS: Component 2 — Paragraphs 3 and 6 both have style 'Heading 2' (0.3 pts)")
            total_score += 0.3
        else:
            if not p3_ok:
                print(f"FAIL: Component 2 — Paragraph 3 expected 'Heading 2', found '{style3}'")
            if not p6_ok:
                print(f"FAIL: Component 2 — Paragraph 6 expected 'Heading 2', found '{style6}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body paragraphs (2,4,5,7,8) have first-line indent ~0.75 cm (0.3 points)
    # 0.75 cm = 270000 EMU; tolerance ±2000 EMU for minor rounding differences
    try:
        body_para_indices = [1, 3, 4, 6, 7]  # 0-based: paragraphs 2,4,5,7,8
        indent_results = []
        for idx in body_para_indices:
            para = doc.paragraphs[idx]
            fli = para.paragraph_format.first_line_indent
            para_num = idx + 1
            if fli is None:
                print(f"FAIL: Component 3 — Paragraph {para_num} has no first-line indent (expected ~0.75 cm)")
                indent_results.append(False)
            else:
                diff = abs(fli - TARGET_INDENT_EMU)
                if diff <= INDENT_TOLERANCE_EMU:
                    cm_val = round(fli / 914400 * 2.54, 4)
                    print(f"  PASS: Paragraph {para_num} first-line indent = {cm_val} cm (EMU={fli})")
                    indent_results.append(True)
                else:
                    cm_val = round(fli / 914400 * 2.54, 4)
                    print(f"  FAIL: Paragraph {para_num} first-line indent = {cm_val} cm (EMU={fli}), expected ~0.75 cm (EMU={TARGET_INDENT_EMU}±{INDENT_TOLERANCE_EMU})")
                    indent_results.append(False)

        if all(indent_results):
            print(f"PASS: Component 3 — All 5 body paragraphs have first-line indent ~0.75 cm (0.3 pts)")
            total_score += 0.3
        else:
            failed = sum(1 for r in indent_results if not r)
            print(f"FAIL: Component 3 — {failed}/5 body paragraphs do not have correct first-line indent")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.1f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
