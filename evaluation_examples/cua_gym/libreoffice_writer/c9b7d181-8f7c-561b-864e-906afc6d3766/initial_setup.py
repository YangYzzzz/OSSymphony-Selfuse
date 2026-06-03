"""
Initial Setup: Information Security Policy document
Task ID: writer_para_055
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm

WORKDIR = '/home/user'
TASK_ID = 'writer_para_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title - currently Default Paragraph Style (NOT Heading 1)
    p1 = doc.add_paragraph('Information Security Policy')
    # Keep as default style (no heading style applied)

    # Paragraph 2: Effective date - body paragraph, no first-line indent yet
    p2 = doc.add_paragraph('Effective Date: March 1, 2025 | Version 3.1')

    # Paragraph 3: Section heading - currently Default Paragraph Style (NOT Heading 2)
    p3 = doc.add_paragraph('Purpose and Scope')

    # Paragraph 4: Body paragraph - no first-line indent yet
    p4 = doc.add_paragraph(
        'This policy establishes the framework for protecting the confidentiality, '
        'integrity, and availability of all information assets owned or managed by the organization.'
    )

    # Paragraph 5: Body paragraph - no first-line indent yet
    p5 = doc.add_paragraph(
        'All employees, contractors, and third-party partners with access to organizational '
        'systems are required to comply with this policy.'
    )

    # Paragraph 6: Section heading - currently Default Paragraph Style (NOT Heading 2)
    p6 = doc.add_paragraph('Password Requirements')

    # Paragraph 7: Body paragraph - no first-line indent yet
    p7 = doc.add_paragraph(
        'All passwords must be a minimum of 12 characters and include at least one uppercase '
        'letter, one lowercase letter, one digit, and one special character.'
    )

    # Paragraph 8: Body paragraph - no first-line indent yet
    p8 = doc.add_paragraph(
        'Passwords must be changed every 90 days. The system will enforce a history of the '
        'last 12 passwords to prevent reuse.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
