"""
Initial Setup: Design document with a rectangle shape with thin solid black border on page 1.
Task ID: writer_obj_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_010'
OUTPUT = f'{WORKDIR}/design_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# Rectangle shape paragraph XML — solid black 1pt border
SHAPE_XML = '''\
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="2743200" cy="1371600"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="1" name="Rectangle 1"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
            <wps:wsp>
              <wps:cNvSpPr>
                <a:spLocks noChangeArrowheads="1"/>
              </wps:cNvSpPr>
              <wps:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="2743200" cy="1371600"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
                <a:noFill/>
                <a:ln w="12700">
                  <a:solidFill>
                    <a:srgbClr val="000000"/>
                  </a:solidFill>
                  <a:prstDash val="solid"/>
                </a:ln>
              </wps:spPr>
              <wps:bodyPr/>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>'''


def create_initial():
    """Create the initial design_doc.docx with a rectangle shape on page 1."""
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add a heading
    doc.add_heading('Project Design Document', level=1)

    # Add introductory paragraph
    doc.add_paragraph(
        'This document outlines the design specifications for the Q2 2025 product launch. '
        'The following sections describe interface layouts, component responsibilities, and delivery timelines.'
    )

    # Add a paragraph explaining the shape below
    label_para = doc.add_paragraph('Figure 1: System Architecture Boundary')
    label_para.paragraph_format.space_before = Pt(6)
    label_para.paragraph_format.space_after = Pt(6)

    # Insert rectangle shape with black solid 1pt border
    shape_elem = etree.fromstring(SHAPE_XML)
    doc.element.body.append(shape_elem)

    # Add caption below the shape
    caption = doc.add_paragraph('The rectangle above represents the core system boundary for Phase 1 deployment.')
    caption.paragraph_format.space_before = Pt(6)

    # Add more content
    doc.add_heading('Key Objectives', level=2)
    objectives = [
        'Deliver a minimum viable product by Q2 end.',
        'Integrate with existing authentication infrastructure.',
        'Support at least 10,000 concurrent users on initial deployment.',
        'Achieve 99.5% uptime SLA for production services.',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('Technical Stack', level=2)
    doc.add_paragraph(
        'The architecture uses a microservices approach with containerized workloads. '
        'Frontend is built on React 18, while the backend exposes REST and GraphQL APIs. '
        'Data persistence is handled via PostgreSQL with Redis caching.'
    )

    # Team info table
    doc.add_heading('Team Responsibilities', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Team'
    hdr_cells[1].text = 'Lead'
    hdr_cells[2].text = 'Deliverable'
    for hdr_cell in hdr_cells:
        run = hdr_cell.paragraphs[0].runs[0]
        run.bold = True

    team_data = [
        ('Frontend', 'Aisha Okonkwo', 'React UI components'),
        ('Backend', 'Lars Eriksson', 'API services and database'),
        ('DevOps', 'Priya Nair', 'CI/CD pipeline and infrastructure'),
        ('QA', 'Tomas Herrera', 'Test automation and coverage'),
    ]
    for team, lead, deliverable in team_data:
        row_cells = table.add_row().cells
        row_cells[0].text = team
        row_cells[1].text = lead
        row_cells[2].text = deliverable

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
