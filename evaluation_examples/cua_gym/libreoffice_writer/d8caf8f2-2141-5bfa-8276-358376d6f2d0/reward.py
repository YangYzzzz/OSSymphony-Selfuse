"""
Reward Script: Color-code every word based on vowel/consonant initial
Task ID: osworld_writer_vowel_consonant_coloring_009
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.4): All vowel-initial words are colored red (FF0000)
  - Component 2 (0.4): All consonant-initial words are colored blue (0000FF)
  - Component 3 (0.2): Title is also color-coded (at least one red + one blue word in para 0)
"""

import os
from docx import Document
from docx.shared import RGBColor

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_009'

RED = RGBColor(0xFF, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xFF)
VOWELS = set('AEIOU')


def color_distance(c1, c2):
    """Euclidean distance between two RGBColor tuples."""
    return ((c1[0] - c2[0]) ** 2 + (c1[1] - c2[1]) ** 2 + (c1[2] - c2[2]) ** 2) ** 0.5


def is_red(color):
    """Check if color is approximately red (FF0000)."""
    if color is None:
        return False
    return color_distance((color[0], color[1], color[2]), (255, 0, 0)) < 50


def is_blue(color):
    """Check if color is approximately blue (0000FF)."""
    if color is None:
        return False
    return color_distance((color[0], color[1], color[2]), (0, 0, 255)) < 50


def starts_with_vowel(word):
    """Return True if word (stripped of leading punctuation) starts with a vowel."""
    clean = word.strip()
    if not clean:
        return False
    return clean[0].upper() in VOWELS


def verify_task(file_path):
    """
    Verify task completion: every word in the document is color-coded
    red (vowel-initial) or blue (consonant-initial).
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Precondition: check that the document has paragraphs with text
    paragraphs_with_text = [p for p in doc.paragraphs if p.text.strip()]
    if not paragraphs_with_text:
        print("CRITICAL: Document has no text paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # Collect all word-runs across all paragraphs
    all_word_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                all_word_runs.append((para, run))

    # Precondition: task requires that words are split into individual runs (colored).
    # Initial env has ~4 runs (one per paragraph). If very few runs, coloring was not applied.
    if len(all_word_runs) < 20:
        print("FAIL: Very few word runs found (%d). Coloring does not appear to have been applied." % len(all_word_runs))
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: All vowel-initial words are colored red (FF0000) — 0.4 pts
    # -----------------------------------------------------------------------
    vowel_words = []
    vowel_correct = 0
    vowel_wrong = []

    try:
        for para, run in all_word_runs:
            word = run.text.strip()
            if not word:
                continue
            if starts_with_vowel(word):
                vowel_words.append((word, run))
                color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                if is_red(color):
                    vowel_correct += 1
                else:
                    vowel_wrong.append((word, color))

        if not vowel_words:
            print("FAIL: Component 1 — No vowel-initial words found in document")
        elif len(vowel_wrong) == 0:
            print("PASS: Component 1 — All %d vowel-initial words colored red (0.4 pts)" % len(vowel_words))
            total_score += 0.4
        else:
            # Partial credit based on proportion correct
            ratio = vowel_correct / len(vowel_words)
            partial = round(0.4 * ratio, 2)
            print("PARTIAL: Component 1 — %d/%d vowel-initial words correctly colored red (%.2f/0.4 pts)" % (
                vowel_correct, len(vowel_words), partial))
            print("  First few wrong: %s" % str(vowel_wrong[:5]))
            if partial > 0:
                total_score += partial
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # -----------------------------------------------------------------------
    # Component 2: All consonant-initial words colored blue (0000FF) — 0.4 pts
    # -----------------------------------------------------------------------
    consonant_words = []
    consonant_correct = 0
    consonant_wrong = []

    try:
        for para, run in all_word_runs:
            word = run.text.strip()
            if not word:
                continue
            if not starts_with_vowel(word):
                consonant_words.append((word, run))
                color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                if is_blue(color):
                    consonant_correct += 1
                else:
                    consonant_wrong.append((word, color))

        if not consonant_words:
            print("FAIL: Component 2 — No consonant-initial words found in document")
        elif len(consonant_wrong) == 0:
            print("PASS: Component 2 — All %d consonant-initial words colored blue (0.4 pts)" % len(consonant_words))
            total_score += 0.4
        else:
            ratio = consonant_correct / len(consonant_words)
            partial = round(0.4 * ratio, 2)
            print("PARTIAL: Component 2 — %d/%d consonant-initial words correctly colored blue (%.2f/0.4 pts)" % (
                consonant_correct, len(consonant_words), partial))
            print("  First few wrong: %s" % str(consonant_wrong[:5]))
            if partial > 0:
                total_score += partial
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # -----------------------------------------------------------------------
    # Component 3: Title paragraph (para 0) is also color-coded — 0.2 pts
    # The title "Echoes of an Empty Afternoon" has all vowel-initial words,
    # so all title runs should be red. We verify that title words are
    # correctly colored according to the vowel/consonant rule.
    # -----------------------------------------------------------------------
    try:
        title_para = doc.paragraphs[0] if doc.paragraphs else None
        if title_para is None or not title_para.text.strip():
            print("FAIL: Component 3 — Title paragraph not found")
        else:
            title_runs = [r for r in title_para.runs if r.text.strip()]
            if not title_runs:
                print("FAIL: Component 3 — Title paragraph has no word runs")
            else:
                title_correct = 0
                title_wrong = []
                for r in title_runs:
                    word = r.text.strip()
                    color = r.font.color.rgb if r.font.color and r.font.color.rgb else None
                    if starts_with_vowel(word):
                        if is_red(color):
                            title_correct += 1
                        else:
                            title_wrong.append((word, 'expected RED, got %s' % color))
                    else:
                        if is_blue(color):
                            title_correct += 1
                        else:
                            title_wrong.append((word, 'expected BLUE, got %s' % color))

                if len(title_wrong) == 0:
                    print("PASS: Component 3 — All %d title words correctly colored (0.2 pts)" % len(title_runs))
                    total_score += 0.2
                elif title_correct > 0:
                    ratio = title_correct / len(title_runs)
                    partial = round(0.2 * ratio, 2)
                    print("PARTIAL: Component 3 — %d/%d title words correctly colored (%.2f/0.2 pts)" % (
                        title_correct, len(title_runs), partial))
                    print("  Wrong: %s" % str(title_wrong[:5]))
                    if partial > 0:
                        total_score += partial
                else:
                    print("FAIL: Component 3 — Title not color-coded")
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    final_score = min(total_score, 1.0)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Default: test against canonical artifact path in VM environment
file_path = '%s/%s.docx' % (WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
