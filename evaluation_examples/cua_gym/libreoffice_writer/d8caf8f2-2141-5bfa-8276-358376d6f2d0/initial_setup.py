"""
Initial Setup: Vowel/Consonant color-coding task
Task ID: osworld_writer_vowel_consonant_coloring_009
Domain: libreoffice_writer

Creates a short story document (~180 words) with title + 3 paragraphs.
All text is in default black — no color coding applied.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Echoes of an Empty Afternoon")
    title_run.bold = True
    title_run.font.size = Pt(16)
    # Default black color (no explicit color set)

    # Paragraph 1
    p1 = doc.add_paragraph()
    p1_run = p1.add_run(
        "Once upon a quiet Tuesday, Oliver wandered into the garden behind his "
        "old stone cottage. Autumn leaves carpeted every inch of the worn flagstone "
        "path, and a gentle breeze carried the faint scent of apples from the "
        "orchard across the lane."
    )
    # Default black

    # Paragraph 2
    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        "Under the ancient oak, a small wooden bench offered a perfect resting "
        "spot. Oliver sat down, placed his notebook on his knee, and began to "
        "sketch the outline of the crumbling wall beyond the rose bushes. Every "
        "scratch of his pencil felt unhurried and alive."
    )
    # Default black

    # Paragraph 3
    p3 = doc.add_paragraph()
    p3_run = p3.add_run(
        "Evening arrived before he noticed the fading light. He closed his "
        "notebook, tucked it under his arm, and stepped inside. On the kitchen "
        "table sat a letter addressed in unfamiliar handwriting. Oliver set down "
        "his pencil case and opened the envelope with careful, trembling fingers."
    )
    # Default black

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
