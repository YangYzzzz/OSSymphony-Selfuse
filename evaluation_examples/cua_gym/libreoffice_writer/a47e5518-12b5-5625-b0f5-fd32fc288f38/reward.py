"""
FINAL REWARD SCRIPT - SUCCESS
Task: Just imported a 30-page OCR scan and it’s riddled with double spaces—literally every sentence has "  " where there should be " ". In LibreOffice Writer, is there a one-shot way to sweep through the whole document and replace every occurrence of two consecutive spaces ('  ') with a single space (' ') instead of hunting them down manually?
Generated: 2025-09-10 13:33:18
Status: success
Model: azure-o3
Total Steps: 1
"""

import os
import re
from docx import Document


def _count_double_spaces(text: str) -> int:
    """Return the number of occurrences of two-or-more consecutive spaces."""
    # Two **or more** spaces in a row – we use `{2,}` so that "   " also counts
    return len(re.findall(r" {2,}", text))


def _gather_all_text(doc: Document) -> str:
    """Concatenate **all** textual content from the DOCX into one string."""
    parts = []

    # 1) Body paragraphs
    for p in doc.paragraphs:
        parts.append(p.text)

    # 2) Table text (tables can appear in OCR-converted docs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    # 3) Headers & footers (rare, but collect to be thorough)
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                parts.append(p.text)
        if section.footer:
            for p in section.footer.paragraphs:
                parts.append(p.text)

    return "\n".join(parts)


def verify_double_space_cleanup(file_path: str) -> float:
    """Reward script for the LibreOffice Writer task.

    Task: Ensure **all** double-space occurrences have been replaced with single spaces.
    A perfect clean-up (0 double spaces) ⇒ reward 1.0.
    Partial credit awarded if only a handful remain; zero credit if the problem largely persists.
    """

    print(f"Verifying double-space removal in: {file_path}")

    # ---------- Prerequisite checks (no points awarded here) ----------
    if not os.path.exists(file_path):
        print("✗ File not found – cannot verify task")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Could not open DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- Actual verification logic ----------
    all_text = _gather_all_text(doc)
    total_spaces = all_text.count(" ")
    double_space_count = _count_double_spaces(all_text)

    print(f"Total spaces found: {total_spaces}")
    print(f"Double-space occurrences: {double_space_count}")

    # Progressive scoring based on remaining double spaces
    score = 0.0
    if double_space_count == 0:
        score = 1.0
        print("✓ No double spaces detected – perfect completion (1.0)")
    elif double_space_count <= 5:
        score = 0.8
        print("✓ Only a few (≤5) double spaces remain – near perfect (0.8)")
    elif double_space_count <= 20:
        score = 0.5
        print("✧ Some (≤20) double spaces remain – partial completion (0.5)")
    elif double_space_count <= 50:
        score = 0.2
        print("✧ Many (≤50) double spaces remain – limited completion (0.2)")
    else:
        score = 0.0
        print("✗ Too many double spaces (>50) – task not completed (0.0)")

    # Ensure the score is a float between 0 and 1
    score = float(min(max(score, 0.0), 1.0))
    print(f"REWARD: {score}")
    return score


# ----------------- Script entry point -----------------
if __name__ == "__main__":
    # Path provided by the task context
    FILE_PATH = "/home/user/just_imported_a_30_page_ocr_scan_and_its_riddled_with_double_spacesliterally_every_sentence_has_wher.docx"
    verify_double_space_cleanup(FILE_PATH)

