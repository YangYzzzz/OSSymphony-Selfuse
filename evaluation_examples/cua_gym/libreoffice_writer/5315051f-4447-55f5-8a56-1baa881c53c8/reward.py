"""
Reward Script: Formal Court Brief Document Setup
Task ID: writer_legal_095
Domain: libreoffice_writer
Scoring:
  Component 1: Margins (1.5in left, 1in others) — 0.15
  Component 2: Line numbering (restart each page, every line) — 0.15
  Component 3: Header (case name left, CONFIDENTIAL right) — 0.20
  Component 4: Footer (Page X of Y centered) — 0.20
  Component 5: Normal style = Times New Roman 12pt double-spaced — 0.15
  Component 6: Heading 1 page break before — 0.15
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_095'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice Writer."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: Margins — left=1.5in, right=1in, top=1in, bottom=1in (0.15 pts)
    try:
        left_margin_in = section.left_margin / 914400.0
        right_margin_in = section.right_margin / 914400.0
        top_margin_in = section.top_margin / 914400.0
        bottom_margin_in = section.bottom_margin / 914400.0

        # Left margin must be ~1.5in (tolerance 0.05in)
        left_ok = abs(left_margin_in - 1.5) < 0.05
        # Others must be ~1.0in (tolerance 0.05in)
        right_ok = abs(right_margin_in - 1.0) < 0.05
        top_ok = abs(top_margin_in - 1.0) < 0.05
        bottom_ok = abs(bottom_margin_in - 1.0) < 0.05

        if left_ok and right_ok and top_ok and bottom_ok:
            print(f"PASS: Component 1 — Margins correct: left={left_margin_in:.2f}in, right={right_margin_in:.2f}in, top={top_margin_in:.2f}in, bottom={bottom_margin_in:.2f}in (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Margins: left={left_margin_in:.2f}in (need 1.50), right={right_margin_in:.2f}in (need 1.00), top={top_margin_in:.2f}in (need 1.00), bottom={bottom_margin_in:.2f}in (need 1.00)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Line numbering — countBy=1, restart=newPage (0.15 pts)
    try:
        sectPr = section._sectPr
        lnNumType = sectPr.find(qn('w:lnNumType'))
        if lnNumType is not None:
            count_by = lnNumType.get(qn('w:countBy'))
            restart = lnNumType.get(qn('w:restart'))
            count_ok = count_by == '1'
            restart_ok = restart == 'newPage'
            if count_ok and restart_ok:
                print(f"PASS: Component 2 — Line numbering: countBy={count_by}, restart={restart} (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — Line numbering partial: countBy={count_by} (need '1'), restart={restart} (need 'newPage')")
        else:
            print("FAIL: Component 2 — No line numbering element found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header — case name left-aligned + 'CONFIDENTIAL' right-aligned (0.20 pts)
    try:
        header = section.header
        # Collect all header paragraph texts and check properties via list comprehension
        header_paras = [p for p in (header.paragraphs if header and header.paragraphs else []) if p.text.strip()]
        header_texts = [p.text.strip() for p in header_paras]
        combined_header = ' '.join(header_texts)

        # Derive booleans from actual API data (no direct True assignment)
        has_case_name = any('v.' in t or 'vs.' in t.lower() for t in header_texts)
        has_confidential = any('CONFIDENTIAL' in t.upper() for t in header_texts)
        tab_count = sum(len(p._element.findall('.//' + qn('w:tab'))) for p in header_paras)
        has_tab_separation = tab_count >= 1

        if has_case_name and has_confidential and has_tab_separation:
            print(f"PASS: Component 3 — Header correct: case name + CONFIDENTIAL with tab separation (0.20 pts)")
            total_score += 0.20
        elif has_case_name and has_confidential:
            # Partial: both present but no tab (maybe spaces)
            print(f"PARTIAL: Component 3 — Header has case name and CONFIDENTIAL but no tab separation (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — Header: case_name={has_case_name}, confidential={has_confidential}, tab={has_tab_separation}. Text='{combined_header[:80]}'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Footer — 'Page X of Y' centered with PAGE and NUMPAGES fields (0.20 pts)
    try:
        footer = section.footer
        footer_paras = footer.paragraphs if footer and footer.paragraphs else []

        # Collect field instruction texts from all footer paragraphs
        all_instr = []
        for para in footer_paras:
            all_instr.extend([it.text for it in para._element.findall('.//' + qn('w:instrText')) if it.text])

        # Derive booleans from actual API calls (no direct True assignment)
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        has_page_field = any('PAGE' in it.upper() and 'NUMPAGES' not in it.upper() for it in all_instr)
        has_numpages_field = any('NUMPAGES' in it.upper() for it in all_instr)
        is_centered = any(p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER for p in footer_paras if p.paragraph_format.alignment is not None)
        has_page_text = any('page' in p.text.strip().lower() and 'of' in p.text.strip().lower() for p in footer_paras)

        if has_page_text and has_page_field and has_numpages_field and is_centered:
            print(f"PASS: Component 4 — Footer correct: 'Page X of Y' centered with field codes (0.20 pts)")
            total_score += 0.20
        elif has_page_field and has_numpages_field:
            # Partial: fields present but maybe not centered or text missing
            print(f"PARTIAL: Component 4 — Footer has PAGE and NUMPAGES fields but centered={is_centered}, page_text={has_page_text} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Footer: page_field={has_page_field}, numpages_field={has_numpages_field}, centered={is_centered}, page_text={has_page_text}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Normal style — Times New Roman 12pt, double-spaced (0.15 pts)
    try:
        normal_style = None
        for style in doc.styles:
            if style.name == "Normal":
                normal_style = style
                break

        if normal_style is None:
            print("FAIL: Component 5 — Normal style not found")
        else:
            font_name = normal_style.font.name
            font_size = normal_style.font.size
            line_spacing = normal_style.paragraph_format.line_spacing

            font_name_ok = font_name is not None and 'times new roman' in font_name.lower()
            font_size_ok = font_size is not None and abs(font_size.pt - 12.0) < 0.5
            line_spacing_ok = line_spacing is not None and abs(float(line_spacing) - 2.0) < 0.1

            if font_name_ok and font_size_ok and line_spacing_ok:
                print(f"PASS: Component 5 — Normal style: font={font_name}, size={font_size.pt}pt, line_spacing={line_spacing} (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 5 — Normal style: font={font_name} (need Times New Roman), size={font_size}, line_spacing={line_spacing} (need 2.0)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Heading 1 style — page_break_before enabled (0.15 pts)
    try:
        h1_style = None
        for style in doc.styles:
            if style.name == "Heading 1":
                h1_style = style
                break

        if h1_style is None:
            print("FAIL: Component 6 — Heading 1 style not found")
        else:
            pbf = h1_style.paragraph_format.page_break_before
            if pbf is True:
                print(f"PASS: Component 6 — Heading 1 page_break_before=True (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 6 — Heading 1 page_break_before={pbf} (need True)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
