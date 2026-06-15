"""
Initial Setup: Court brief with default formatting (no legal formatting applied)
Task ID: writer_legal_095
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_095'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Default margins (1 inch all around — python-docx default)
    # No line numbering, no headers/footers
    # Liberation Serif is the default LO font — we leave font at default (Calibri in docx)
    # Single spacing (default)

    # ---- Title Page ----
    title = doc.add_heading('BRIEF IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacer

    caption = doc.add_paragraph()
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption.add_run('IN THE UNITED STATES DISTRICT COURT\n'
                          'FOR THE NORTHERN DISTRICT OF CALIFORNIA\n\n'
                          'MERIDIAN TECHNOLOGIES, INC.,\nPlaintiff,\n\n'
                          'v.\n\n'
                          'ATLAS GLOBAL SOLUTIONS, LLC,\nDefendant.\n\n'
                          'Case No. 3:2025-cv-04817-RMW')
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(12)

    doc.add_page_break()

    # ---- Table of Contents ----
    toc_heading = doc.add_heading('TABLE OF CONTENTS', level=1)

    toc_items = [
        ('I.', 'INTRODUCTION', '1'),
        ('II.', 'STATEMENT OF FACTS', '2'),
        ('III.', 'LEGAL STANDARD', '4'),
        ('IV.', 'ARGUMENT', '5'),
        ('', 'A. Defendant Breached the Master Services Agreement', '5'),
        ('', 'B. Plaintiff Suffered Quantifiable Damages', '8'),
        ('', 'C. No Genuine Dispute of Material Fact Exists', '10'),
        ('V.', 'CONCLUSION', '12'),
    ]
    for num, title_text, page in toc_items:
        p = doc.add_paragraph()
        if num:
            run = p.add_run(f'{num}\t{title_text}')
        else:
            run = p.add_run(f'\t{title_text}')
        run.font.name = 'Liberation Serif'
        run.font.size = Pt(12)
        tab_run = p.add_run(f'\t{page}')
        tab_run.font.name = 'Liberation Serif'
        tab_run.font.size = Pt(12)

    doc.add_page_break()

    # ---- I. INTRODUCTION ----
    doc.add_heading('I. INTRODUCTION', level=1)

    intro_paras = [
        'Plaintiff Meridian Technologies, Inc. ("Meridian") respectfully submits this brief '
        'in support of its Motion for Summary Judgment against Defendant Atlas Global Solutions, '
        'LLC ("Atlas"). As demonstrated below, there is no genuine dispute of material fact, and '
        'Meridian is entitled to judgment as a matter of law on its claims for breach of contract '
        'and unjust enrichment.',

        'On March 15, 2024, Meridian and Atlas entered into a Master Services Agreement '
        '("MSA") pursuant to which Atlas agreed to provide enterprise cloud migration services '
        'for Meridian\'s North American operations. The total contract value was $4,275,000, '
        'payable in quarterly installments tied to specific performance milestones.',

        'Despite receiving the first two quarterly payments totaling $2,137,500, Atlas failed '
        'to complete any of the agreed-upon milestones by the contractual deadlines. Atlas\'s '
        'project manager, Rebecca Torres, acknowledged in a June 2024 email that "the timeline '
        'has slipped significantly and we are unable to meet the Q3 deliverables as specified."',
    ]
    for text in intro_paras:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Liberation Serif'
            run.font.size = Pt(12)

    # ---- II. STATEMENT OF FACTS ----
    doc.add_heading('II. STATEMENT OF FACTS', level=1)

    facts_paras = [
        'Meridian is a Delaware corporation with its principal place of business in San Jose, '
        'California. Founded in 2011, Meridian develops and markets enterprise resource planning '
        'software used by over 3,200 organizations worldwide. In fiscal year 2024, Meridian '
        'reported revenues of $892 million and employed approximately 4,100 people.',

        'Atlas is a limited liability company organized under the laws of Texas, with offices '
        'in Austin, Dallas, and Atlanta. Atlas holds itself out as a specialist in enterprise '
        'cloud migration, data center consolidation, and IT infrastructure modernization.',

        'In early 2024, Meridian issued a Request for Proposals ("RFP") seeking a vendor to '
        'migrate its legacy on-premises infrastructure to a hybrid cloud environment. The RFP '
        'specified a 12-month implementation timeline with four quarterly milestones. Atlas '
        'submitted a proposal on February 8, 2024, which Meridian accepted after a competitive '
        'evaluation process involving five qualified vendors.',

        'The MSA, executed on March 15, 2024, contained the following key provisions relevant '
        'to this motion:',

        '(a) Section 3.1 required Atlas to complete Phase 1 (environment assessment and '
        'architecture design) by June 15, 2024;',

        '(b) Section 3.2 required Atlas to complete Phase 2 (data migration pilot for the '
        'Western Region) by September 15, 2024;',

        '(c) Section 7.4 provided that failure to meet any milestone by more than 30 days '
        'constituted a material breach entitling Meridian to terminate and recover all payments '
        'made;',

        '(d) Section 12.1 contained a limitation of liability capping consequential damages '
        'at 150% of the total contract value.',

        'Atlas assigned a team of 14 consultants to the project, led by Senior Director '
        'Rebecca Torres. Internal Atlas communications produced during discovery reveal that '
        'by April 2024, Atlas had diverted six of these consultants to a competing engagement '
        'with Vanguard Financial Group without notifying Meridian.',
    ]
    for text in facts_paras:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Liberation Serif'
            run.font.size = Pt(12)

    # ---- III. LEGAL STANDARD ----
    doc.add_heading('III. LEGAL STANDARD', level=1)

    legal_paras = [
        'Summary judgment is appropriate when "the movant shows that there is no genuine '
        'dispute as to any material fact and the movant is entitled to judgment as a matter '
        'of law." Fed. R. Civ. P. 56(a). A fact is material if it "might affect the outcome '
        'of the suit under the governing law." Anderson v. Liberty Lobby, Inc., 477 U.S. 242, '
        '248 (1986).',

        'The moving party bears the initial burden of demonstrating the absence of a genuine '
        'issue of material fact. Celotex Corp. v. Catrett, 477 U.S. 317, 323 (1986). Once '
        'this burden is met, the nonmoving party must go beyond the pleadings and "set forth '
        'specific facts showing that there is a genuine issue for trial." Id. at 324.',

        'In the Ninth Circuit, courts must view the evidence in the light most favorable to '
        'the nonmoving party. T.W. Elec. Serv., Inc. v. Pac. Elec. Contractors Ass\'n, 809 '
        'F.2d 626, 630-31 (9th Cir. 1987). However, the nonmoving party "must do more than '
        'simply show that there is some metaphysical doubt as to the material facts." '
        'Matsushita Elec. Indus. Co. v. Zenith Radio Corp., 475 U.S. 574, 586 (1986).',
    ]
    for text in legal_paras:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Liberation Serif'
            run.font.size = Pt(12)

    # ---- IV. ARGUMENT ----
    doc.add_heading('IV. ARGUMENT', level=1)

    arg_paras = [
        'A. Defendant Breached the Master Services Agreement',
        '',
        'Under California law, the elements of breach of contract are: (1) the existence of a '
        'contract; (2) the plaintiff\'s performance or excuse for nonperformance; (3) the '
        'defendant\'s breach; and (4) resulting damages. Oasis West Realty, LLC v. Goldman, '
        '51 Cal. 4th 811, 821 (2011).',

        'Here, all four elements are established by undisputed facts. First, the MSA is a '
        'valid, enforceable contract executed by authorized representatives of both parties. '
        'Atlas does not dispute the existence or validity of the MSA.',

        'Second, Meridian performed all of its obligations under the MSA, including timely '
        'payment of the first two quarterly installments totaling $2,137,500 and providing '
        'Atlas with full access to its IT infrastructure, personnel, and documentation as '
        'required by Section 4.2.',

        'Third, Atlas materially breached the MSA by failing to complete Phase 1 by the '
        'June 15, 2024 deadline, and Phase 2 by the September 15, 2024 deadline. As of the '
        'date of termination (November 3, 2024), Atlas had completed only 23% of the Phase 1 '
        'deliverables — well short of the 100% completion required under Section 3.1.',

        'B. Plaintiff Suffered Quantifiable Damages',
        '',
        'Meridian has suffered direct damages of $2,137,500 in payments made to Atlas for '
        'services never rendered. Additionally, Meridian retained Pinnacle Systems, Inc. to '
        'complete the migration at a cost of $3,850,000 — a premium of $1,712,500 over the '
        'remaining balance of the Atlas contract. Meridian further incurred $425,000 in '
        'internal costs for project delays and $187,500 in legal fees related to the '
        'termination process.',

        'C. No Genuine Dispute of Material Fact Exists',
        '',
        'Atlas cannot point to any genuine dispute regarding the material facts. The MSA '
        'speaks for itself regarding the deadlines and payment terms. Atlas\'s own internal '
        'communications confirm the diversion of resources and failure to meet milestones. '
        'The Torres email of June 2024 is a party admission under Federal Rule of Evidence '
        '801(d)(2)(D).',
    ]
    for text in arg_paras:
        if text.startswith('A. ') or text.startswith('B. ') or text.startswith('C. '):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Liberation Serif'
            run.font.size = Pt(12)
        elif text == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Liberation Serif'
                run.font.size = Pt(12)

    # ---- V. CONCLUSION ----
    doc.add_heading('V. CONCLUSION', level=1)

    conclusion_paras = [
        'For the foregoing reasons, Meridian Technologies, Inc. respectfully requests that '
        'this Court grant its Motion for Summary Judgment and enter judgment in favor of '
        'Meridian on its claims for breach of contract and unjust enrichment, awarding damages '
        'in the amount of $4,462,500, plus pre-judgment interest, costs, and such further '
        'relief as the Court deems just and proper.',

        'Respectfully submitted,',
        '',
        'Jonathan R. Whitfield, Esq.',
        'WHITFIELD & ASSOCIATES LLP',
        '555 California Street, Suite 3200',
        'San Francisco, CA 94104',
        'Tel: (415) 555-0192',
        'Counsel for Plaintiff Meridian Technologies, Inc.',
        '',
        'Dated: January 15, 2026',
    ]
    for text in conclusion_paras:
        if text == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Liberation Serif'
                run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
