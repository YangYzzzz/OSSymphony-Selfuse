"""
Reward Script: Format Revenue column numbers with thousand separators and 2 decimal places
Task ID: writer_tbl_056
Domain: libreoffice_writer
Scoring:
  Component 1: Each of the 4 Revenue data cells matches expected 'N,NNN.00' format (0.175 pts each = 0.7 pts total)
  Component 2: ALL 4 Revenue cells correctly formatted AND other columns (Header, Quarter, Growth) remain unchanged (0.3 pts)
"""

import os
import re

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_056'

# Expected formatted values for column 2 (Revenue), data rows 1-4 (0-indexed)
EXPECTED_REVENUE = ['45,000.00', '52,000.00', '48,000.00', '61,000.00']

# Expected header row
EXPECTED_HEADER = ['Quarter', 'Revenue', 'Growth']

# Expected data for non-revenue columns
EXPECTED_COL0 = ['Q1', 'Q2', 'Q3', 'Q4']
EXPECTED_COL2 = ['5%', '15%', '-8%', '27%']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: Ensure the table exists and has the expected shape
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) != 5 or len(table.columns) != 3:
        print(f"CRITICAL: Table shape mismatch. Expected 5x3, got {len(table.rows)}x{len(table.columns)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Each Revenue data cell is formatted with thousand separators and 2 decimal places
    # 4 cells x 0.175 pts each = 0.7 pts total
    # This FAILS on initial (plain integers) and PASSES on golden (formatted strings)
    correct_cells = 0
    try:
        formatted_pattern = re.compile(r'^\d{1,3}(,\d{3})*\.\d{2}$')
        per_cell_pts = 0.175

        for row_idx in range(1, 5):  # Rows 1-4 (data rows, 0-indexed)
            cell_text = table.cell(row_idx, 1).text.strip()
            expected = EXPECTED_REVENUE[row_idx - 1]
            if cell_text == expected:
                print(f"PASS: Row {row_idx} Revenue = '{cell_text}' (matches expected) (+{per_cell_pts} pts)")
                total_score += per_cell_pts
                correct_cells += 1
            elif formatted_pattern.match(cell_text):
                # Correct format but unexpected value — award half points
                print(f"PARTIAL: Row {row_idx} Revenue = '{cell_text}' — format correct but value differs from expected '{expected}' (+{per_cell_pts * 0.5} pts)")
                total_score += per_cell_pts * 0.5
            else:
                print(f"FAIL: Row {row_idx} Revenue = '{cell_text}' — not formatted (expected '{expected}')")

        comp1_subtotal = round(min(total_score, 0.7), 4)
        print(f"Component 1 subtotal: {comp1_subtotal}/0.7")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: ALL 4 Revenue cells are correctly formatted AND other columns intact (0.3 pts)
    # This is a COMPOUND check — it gates on correct_cells == 4 (all revenue cells match exactly) AND
    # verifies that header/Quarter/Growth columns are unchanged.
    # FAILS on initial (because revenue cells are not formatted), PASSES on golden.
    try:
        if correct_cells == 4:
            comp2_parts = 0

            # Sub-check A: Header row unchanged
            header_ok = all(
                table.cell(0, c).text.strip() == EXPECTED_HEADER[c]
                for c in range(3)
            )
            if header_ok:
                comp2_parts += 1
                print(f"PASS: Header row intact: {EXPECTED_HEADER}")
            else:
                print(f"FAIL: Header row changed or incorrect")

            # Sub-check B: Quarter column (col 0) unchanged
            col0_ok = all(
                table.cell(r, 0).text.strip() == EXPECTED_COL0[r - 1]
                for r in range(1, 5)
            )
            if col0_ok:
                comp2_parts += 1
                print(f"PASS: Quarter column unchanged")
            else:
                print(f"FAIL: Quarter column changed")

            # Sub-check C: Growth column (col 2) unchanged
            col2_ok = all(
                table.cell(r, 2).text.strip() == EXPECTED_COL2[r - 1]
                for r in range(1, 5)
            )
            if col2_ok:
                comp2_parts += 1
                print(f"PASS: Growth column unchanged")
            else:
                print(f"FAIL: Growth column changed")

            # Award 0.3 only if ALL sub-checks pass (all revenue formatted + all columns intact)
            if comp2_parts == 3:
                print(f"PASS: Component 2 — All revenue formatted and other columns intact (+0.3 pts)")
                total_score += 0.3
            elif comp2_parts > 0:
                partial = round((comp2_parts / 3) * 0.3, 4)
                print(f"PARTIAL: Component 2 — {comp2_parts}/3 sub-checks passed (+{partial} pts)")
                total_score += partial

            print(f"Component 2 subtotal: {round((comp2_parts / 3) * 0.3, 4)}/0.3")
        else:
            print(f"FAIL: Component 2 — Skipped because not all Revenue cells are correctly formatted (0.0/0.3)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


file_path = f'{WORKDIR}/revenue_table.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
