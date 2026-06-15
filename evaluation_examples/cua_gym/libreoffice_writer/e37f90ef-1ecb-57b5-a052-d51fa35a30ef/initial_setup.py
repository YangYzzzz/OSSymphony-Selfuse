"""
Initial Setup: Memorial letter to grandfather (pre-formatted state)
Task ID: writer_creative_068
Domain: libreoffice_writer

Creates a heartfelt memorial letter in plain formatting (12pt Times New Roman,
single-spaced, default margins) with [PAUSE] markers at end of paragraphs 2, 4, 6.
File saved to /home/user/Desktop/memorial_letter.docx and opened in LibreOffice Writer.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_068'
OUTPUT = f'{WORKDIR}/memorial_letter.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default margins (standard 1-inch)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # --- Title: "For Grandpa Henry" ---
    # Plain: 12pt Times New Roman, not bold, not centered (left-aligned)
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title_para.add_run("For Grandpa Henry")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(12)
    title_run.font.bold = False

    # --- Subtitle: "Read by his grandson, Daniel Walker" ---
    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    subtitle_run = subtitle_para.add_run("Read by his grandson, Daniel Walker")
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = False

    # --- Date: "March 8, 2026" ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    date_run = date_para.add_run("March 8, 2026")
    date_run.font.name = "Times New Roman"
    date_run.font.size = Pt(12)

    # --- Body paragraphs (6 total; paragraphs 2, 4, 6 end with [PAUSE]) ---
    body_paragraphs = [
        # Paragraph 1 — no [PAUSE]
        (
            "Grandpa Henry, you were the kind of man who could fix anything with a piece of wire "
            "and a patient smile. I remember the long summer afternoons in your garage, the smell "
            "of sawdust and engine oil, and your steady hands guiding mine as you taught me how to "
            "listen to a machine and understand what it needed. You never rushed. You never raised "
            "your voice. You simply believed that everything could be made right with enough care."
        ),
        # Paragraph 2 — ends with [PAUSE]
        (
            "You taught me that real strength is quiet. That a man shows his love not through grand "
            "speeches, but through showing up — through driving three hours in a snowstorm for a "
            "school play, through sitting at the kitchen table long after dinner to help with homework "
            "you didn't entirely understand yourself. You never made it feel like sacrifice. You made "
            "it feel like the most natural thing in the world to give your whole self to the people "
            "you loved. [PAUSE]"
        ),
        # Paragraph 3 — no [PAUSE]
        (
            "I think of you every time I hear old country music drifting from a radio, every time I "
            "see a garden coming back to life in spring, every time I hold a tool in my hand and "
            "feel the weight of it. You are in these ordinary things. You are in the way I try to "
            "stand a little taller when things get hard, in the way I try to be patient when patience "
            "doesn't come easily. You are still teaching me."
        ),
        # Paragraph 4 — ends with [PAUSE]
        (
            "There were so many things I never said out loud while you were here. I never told you "
            "how much I admired your steadiness in hard times, or how many times I've wished I could "
            "call you just to hear your voice say that it's going to be all right. I never told you "
            "that the summer I spent with you and Grandma after my parents divorced was the summer "
            "that saved me. You gave me ground to stand on when everything else felt like it was "
            "shifting. I am grateful beyond words. [PAUSE]"
        ),
        # Paragraph 5 — no [PAUSE]
        (
            "Our family is a reflection of you. The way we gather around the table and tell the same "
            "stories until they become legend. The way we fight hard and forgive faster. The way we "
            "show up for each other without being asked. You built that. Not with speeches or rules, "
            "but with the example of how you lived — quietly, generously, with enormous loyalty and "
            "an unshakeable belief in the goodness of people."
        ),
        # Paragraph 6 — ends with [PAUSE]
        (
            "Grandpa Henry, I promise to carry you forward. I promise to remember what patience looks "
            "like, what grace looks like, what it means to love a family the way you loved ours. "
            "Wherever you are now, I hope you can feel how much you are missed, and how much you are "
            "honored today, not just by the words we say but by the people you helped us become. "
            "Thank you. For everything. For all of it. [PAUSE]"
        ),
    ]

    for text in body_paragraphs:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Single-spaced
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
