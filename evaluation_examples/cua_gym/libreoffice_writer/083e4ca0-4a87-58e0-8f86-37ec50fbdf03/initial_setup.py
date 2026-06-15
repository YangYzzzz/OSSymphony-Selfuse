"""
Initial Setup: Mail merge letter with company database
Task ID: writer_mt_024
Domain: libreoffice_writer

Creates:
1. /home/user/writer_mt_024.docx - Annual review letter template (no merge fields)
2. /home/user/Desktop/company.odb - LibreOffice Base database with Employees, Departments, Projects tables
Opens the docx in LibreOffice Writer.
"""

import os
import shlex
import subprocess
import time
import zipfile
import io
import shutil

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_024'
OUTPUT_DOCX = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP = f'{WORKDIR}/Desktop'
ODB_PATH = f'{DESKTOP}/company.odb'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_odb():
    """Create a LibreOffice Base .odb file with HSQLDB embedded database."""
    os.makedirs(DESKTOP, exist_ok=True)

    # An .odb file is a ZIP (ODF package) containing:
    # - META-INF/manifest.xml
    # - content.xml (database settings)
    # - database/ directory with HSQLDB files

    manifest_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0"
                   manifest:version="1.2">
  <manifest:file-entry manifest:media-type="application/vnd.oasis.opendocument.base"
                       manifest:full-path="/"/>
  <manifest:file-entry manifest:media-type=""
                       manifest:full-path="database/"/>
  <manifest:file-entry manifest:media-type="text/xml"
                       manifest:full-path="content.xml"/>
</manifest:manifest>'''

    content_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
    xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    xmlns:db="urn:oasis:names:tc:opendocument:xmlns:database:1.0"
    office:version="1.2">
  <office:body>
    <office:database>
      <db:data-source>
        <db:connection-data>
          <db:connection-resource db:href="sdbc:embedded:hsqldb"/>
        </db:connection-data>
        <db:login db:is-password-required="false"/>
      </db:data-source>
    </office:database>
  </office:body>
</office:document-content>'''

    # HSQLDB .script file defines schema and data
    hsqldb_script = '''CREATE SCHEMA PUBLIC AUTHORIZATION DBA
CREATE USER SA PASSWORD ""
GRANT DBA TO SA
SET WRITE_DELAY 0
CREATE CACHED TABLE "Employees" ("EmpID" INTEGER NOT NULL PRIMARY KEY, "FirstName" VARCHAR(50), "LastName" VARCHAR(50), "Title" VARCHAR(100), "HireDate" VARCHAR(20), "Salary" DECIMAL(10,2))
CREATE CACHED TABLE "Departments" ("DeptID" INTEGER NOT NULL PRIMARY KEY, "DeptName" VARCHAR(100), "Manager" VARCHAR(100), "Budget" DECIMAL(12,2))
CREATE CACHED TABLE "Projects" ("ProjectID" INTEGER NOT NULL PRIMARY KEY, "ProjectName" VARCHAR(100), "DeptID" INTEGER, "StartDate" VARCHAR(20), "Status" VARCHAR(30))
INSERT INTO "Employees" VALUES(1001,'Sarah','Chen','Senior Software Engineer','2021-03-15',95000.00)
INSERT INTO "Employees" VALUES(1002,'Marcus','Johnson','Marketing Director','2019-08-01',112000.00)
INSERT INTO "Employees" VALUES(1003,'Elena','Rodriguez','HR Manager','2020-01-10',88000.00)
INSERT INTO "Employees" VALUES(1004,'James','Williams','Data Analyst','2022-06-20',76000.00)
INSERT INTO "Employees" VALUES(1005,'Priya','Patel','Product Manager','2020-11-05',98000.00)
INSERT INTO "Employees" VALUES(1006,'David','Kim','DevOps Lead','2021-09-12',91000.00)
INSERT INTO "Employees" VALUES(1007,'Olivia','Thompson','UX Designer','2023-02-28',82000.00)
INSERT INTO "Employees" VALUES(1008,'Robert','Garcia','Finance Controller','2018-04-15',105000.00)
INSERT INTO "Departments" VALUES(101,'Engineering','Sarah Chen',1500000.00)
INSERT INTO "Departments" VALUES(102,'Marketing','Marcus Johnson',800000.00)
INSERT INTO "Departments" VALUES(103,'Human Resources','Elena Rodriguez',450000.00)
INSERT INTO "Departments" VALUES(104,'Finance','Robert Garcia',600000.00)
INSERT INTO "Projects" VALUES(201,'Platform Migration',101,'2025-01-15','In Progress')
INSERT INTO "Projects" VALUES(202,'Brand Refresh',102,'2025-03-01','Planning')
INSERT INTO "Projects" VALUES(203,'Employee Portal',103,'2024-11-01','Active')
INSERT INTO "Projects" VALUES(204,'Budget Automation',104,'2025-02-10','In Progress')
'''

    hsqldb_properties = '''#HSQL Database Engine 1.8.0.10
#Thu Jan 01 00:00:00 UTC 2026
hsqldb.script_format=0
runtime.gc_interval=0
sql.enforce_strict_size=true
hsqldb.cache_size_scale=8
readonly=false
hsqldb.nio_data_file=true
hsqldb.cache_scale=13
version=1.8.0
hsqldb.default_table_type=cached
hsqldb.cache_file_scale=1
hsqldb.log_size=10
modified=no
hsqldb.compatible_version=1.8.0
'''

    with zipfile.ZipFile(ODB_PATH, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('META-INF/manifest.xml', manifest_xml)
        zf.writestr('content.xml', content_xml)
        zf.writestr('database/script', hsqldb_script)
        zf.writestr('database/properties', hsqldb_properties)
        # Create empty data file (needed by HSQLDB)
        zf.writestr('database/data', '')

    print(f'ODB database created: {ODB_PATH}')


def create_docx():
    """Create the annual review letter template without merge fields."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Company header
    header_para = doc.add_paragraph()
    header_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.paragraph_format.space_after = Pt(6)
    run = header_para.add_run('Vertex Dynamics Inc.')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

    sub_header = doc.add_paragraph()
    sub_header.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_header.paragraph_format.space_after = Pt(24)
    run = sub_header.add_run('Annual Performance Review Notification')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    # Date line
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.add_run('March 28, 2026')
    run.font.size = Pt(11)

    # Salutation placeholder area
    sal_para = doc.add_paragraph()
    sal_para.paragraph_format.space_after = Pt(12)
    run = sal_para.add_run('Dear [Employee],')
    run.font.size = Pt(11)

    # Body paragraph 1
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(8)
    run = body1.add_run(
        'We are pleased to inform you that your annual performance review has been scheduled. '
        'As a valued member of our team, your contributions over the past year have been '
        'instrumental in driving the success of Vertex Dynamics Inc.'
    )
    run.font.size = Pt(11)

    # Body paragraph 2 - with placeholder for merge fields
    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(8)
    run = body2.add_run(
        'According to our records, you joined the company on [HireDate] as [Title]. '
        'During your tenure, you have consistently demonstrated excellence in your role.'
    )
    run.font.size = Pt(11)

    # Body paragraph 3
    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(8)
    run = body3.add_run(
        'Your review meeting will be conducted by your direct supervisor during the week '
        'of April 14-18, 2026. Please prepare a brief self-assessment highlighting your '
        'key achievements, challenges overcome, and professional development goals for '
        'the coming year.'
    )
    run.font.size = Pt(11)

    # Body paragraph 4
    body4 = doc.add_paragraph()
    body4.paragraph_format.space_after = Pt(8)
    run = body4.add_run(
        'If you have any questions regarding the review process, please do not hesitate '
        'to contact the Human Resources department at hr@vertexdynamics.com or extension 4500.'
    )
    run.font.size = Pt(11)

    # Closing
    closing1 = doc.add_paragraph()
    closing1.paragraph_format.space_before = Pt(18)
    closing1.paragraph_format.space_after = Pt(4)
    run = closing1.add_run('Sincerely,')
    run.font.size = Pt(11)

    closing2 = doc.add_paragraph()
    closing2.paragraph_format.space_after = Pt(2)
    run = closing2.add_run('Elena Rodriguez')
    run.bold = True
    run.font.size = Pt(11)

    closing3 = doc.add_paragraph()
    run = closing3.add_run('HR Manager, Vertex Dynamics Inc.')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.save(OUTPUT_DOCX)
    print(f'Document created: {OUTPUT_DOCX}')


def main():
    create_odb()
    create_docx()

    # Open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT_DOCX}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
