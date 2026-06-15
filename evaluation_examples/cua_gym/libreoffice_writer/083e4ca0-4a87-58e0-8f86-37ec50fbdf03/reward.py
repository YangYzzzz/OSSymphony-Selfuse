"""
Reward Script: Connect to company.odb and use Employees table as mail merge data source
Task ID: writer_mt_024
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): All 4 required MERGEFIELD instrText elements present (FirstName, LastName, HireDate, Title)
  Component 2 (0.3): Placeholders replaced - original bracket placeholders removed from document text
  Component 3 (0.2): Correct field structure - fldChar elements present (>=12 for 4 complete fields)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_024'

# Required merge field names from the task specification
REQUIRED_FIELDS = {'FirstName', 'LastName', 'HireDate', 'Title'}


def persist_app_state(domain: str):
    """Best-effort save of any unsaved GUI state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify mail merge field insertion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from lxml import etree
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body

    # Component 1: All 4 required MERGEFIELD instrText elements present (0.5 points)
    # This is the primary task requirement - inserting merge fields for FirstName, LastName, HireDate, Title
    try:
        instrs = body.findall('.//w:instrText', ns)
        found_fields = set()
        for inst in instrs:
            if inst.text and 'MERGEFIELD' in inst.text:
                # Extract field name from instrText like ' MERGEFIELD FirstName '
                match = re.search(r'MERGEFIELD\s+(\w+)', inst.text)
                if match:
                    found_fields.add(match.group(1))

        matched_fields = found_fields & REQUIRED_FIELDS
        num_matched = len(matched_fields)
        if num_matched == len(REQUIRED_FIELDS):
            print(f"PASS: Component 1 - All 4 required merge fields found: {matched_fields} (0.5 pts)")
            total_score += 0.5
        elif num_matched > 0:
            # Partial credit: proportional to fields found
            if num_matched >= 1:
                partial = 0.5 * (num_matched / len(REQUIRED_FIELDS))
                print(f"PARTIAL: Component 1 - Found {num_matched}/{len(REQUIRED_FIELDS)} merge fields: {matched_fields} ({partial:.2f} pts)")
                total_score += partial
        else:
            print(f"FAIL: Component 1 - No required merge fields found. instrText elements: {[i.text for i in instrs]}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Placeholders replaced (0.3 points)
    # Initial doc has [Employee], [HireDate], [Title] as plain text.
    # Golden doc should NOT have these bracket placeholders anymore.
    try:
        full_text = '\n'.join(p.text for p in doc.paragraphs)

        # Check that old placeholders are gone
        old_placeholders_present = (
            '[Employee]' in full_text or
            '[HireDate]' in full_text or
            '[Title]' in full_text
        )

        # Check that merge field display text is present (guillemets or field text)
        # When merge fields are inserted, python-docx para.text may show the field result text
        # or the guillemet markers. We check that the instrText fields exist (covered in C1)
        # and that old placeholders are gone.
        # Check that merge field display text is present in paragraph text
        merge_display_count = sum(
            1 for p in doc.paragraphs
            if any(fname in p.text for fname in ['FirstName', 'LastName', 'HireDate', 'Title'])
        )

        if not old_placeholders_present and merge_display_count > 0:
            print(f"PASS: Component 2 - Old placeholders removed and merge field text present (0.3 pts)")
            total_score += 0.3
        elif not old_placeholders_present:
            print(f"PARTIAL: Component 2 - Old placeholders removed but no merge field display text found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 - Old bracket placeholders still present in document text")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Correct field structure with fldChar elements (0.2 points)
    # Each MERGEFIELD needs 3 fldChar elements (begin, separate/end pattern).
    # 4 fields x 3 = 12 minimum fldChar elements.
    try:
        fld_chars = body.findall('.//w:fldChar', ns)
        fld_count = len(fld_chars)

        if fld_count >= 12:
            print(f"PASS: Component 3 - Found {fld_count} fldChar elements (>= 12 for 4 fields) (0.2 pts)")
            total_score += 0.2
        elif fld_count >= 6:
            # At least 2 fields worth of structure
            print(f"PARTIAL: Component 3 - Found {fld_count} fldChar elements (partial field structure) (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 3 - Found only {fld_count} fldChar elements (expected >= 12)")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
