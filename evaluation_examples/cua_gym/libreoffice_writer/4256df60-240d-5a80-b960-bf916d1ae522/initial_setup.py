"""
Initial Setup: Add page numbers in footer with 'Page X of Y' format
Task ID: writer_tm_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_053'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Ensure footer is EMPTY (no page numbers)
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear any default footer content
    for para in footer.paragraphs:
        para.text = ""

    # === PAGE 1 ===
    heading = doc.add_heading('Quarterly Performance Report', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    p = doc.add_paragraph('Prepared by: Strategic Planning Division')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph('Date: March 15, 2025')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph('Classification: Internal Use Only')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')
    doc.add_paragraph(
        'This report provides a comprehensive analysis of our organizational performance '
        'across all major business units for Q4 2024. The findings presented herein are '
        'based on data collected from regional offices, financial systems, and customer '
        'feedback platforms.'
    )

    doc.add_page_break()

    # === PAGE 2 ===
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'The fourth quarter of 2024 demonstrated strong growth across multiple segments. '
        'Total revenue reached $142.3 million, representing a 12.7% increase year-over-year. '
        'Our North American operations contributed $89.5 million, while European markets '
        'generated $34.8 million. The Asia-Pacific region showed the most significant growth '
        'at 23.4%, reaching $18.0 million in quarterly revenue.'
    )
    doc.add_paragraph(
        'Operating expenses were maintained at $98.7 million, reflecting disciplined cost '
        'management across departments. The engineering division invested $22.4 million in '
        'research and development, focusing on next-generation product features and '
        'infrastructure modernization. Marketing expenditures totaled $15.3 million, '
        'with digital campaigns accounting for 67% of the total marketing budget.'
    )
    doc.add_paragraph(
        'Customer satisfaction scores improved by 8.2 points to reach 87.3 out of 100, '
        'driven by enhanced support response times and product reliability improvements. '
        'Net Promoter Score increased from 42 to 51, indicating growing customer loyalty '
        'and brand advocacy across our primary market segments.'
    )
    doc.add_paragraph(
        'Employee headcount grew by 127 positions to a total of 2,847 full-time equivalents. '
        'Key hires included senior leadership in the data science and product management '
        'functions. Voluntary turnover decreased to 8.3%, down from 11.2% in the prior quarter.'
    )

    doc.add_page_break()

    # === PAGE 3 ===
    doc.add_heading('Financial Performance', level=1)

    doc.add_heading('Revenue Breakdown', level=2)
    doc.add_paragraph(
        'Product revenue accounted for $104.7 million (73.6% of total), while services '
        'revenue contributed $37.6 million (26.4%). Subscription-based recurring revenue '
        'grew to $67.2 million, now representing 47.2% of total revenue and reflecting '
        'our successful transition toward predictable revenue streams.'
    )

    doc.add_heading('Regional Analysis', level=2)
    # Add a table for regional data
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Region', 'Revenue ($M)', 'Growth (%)', 'Headcount']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True

    data = [
        ['North America', '$89.5', '10.3%', '1,542'],
        ['Europe', '$34.8', '8.7%', '723'],
        ['Asia-Pacific', '$18.0', '23.4%', '582'],
        ['Total', '$142.3', '12.7%', '2,847'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'The Asia-Pacific expansion strategy yielded exceptional results, with the Singapore '
        'office achieving profitability ahead of schedule. The Tokyo team secured three '
        'enterprise contracts valued at over $2 million each, including partnerships with '
        'Mitsubishi Electric, Sumitomo Corporation, and Rakuten Group.'
    )

    doc.add_page_break()

    # === PAGE 4 ===
    doc.add_heading('Operational Highlights', level=1)

    doc.add_heading('Product Development', level=2)
    doc.add_paragraph(
        'The engineering team delivered 14 major feature releases during Q4, including the '
        'highly anticipated analytics dashboard redesign and the real-time collaboration '
        'module. Platform uptime exceeded 99.97%, surpassing our SLA commitment of 99.9%. '
        'Average API response time improved from 245ms to 178ms following infrastructure '
        'optimization efforts.'
    )

    doc.add_heading('Customer Success', level=2)
    doc.add_paragraph(
        'The customer success team onboarded 234 new enterprise clients during the quarter. '
        'Average time-to-value decreased from 42 days to 28 days through improved onboarding '
        'workflows and dedicated implementation specialists. Contract renewal rates reached '
        '94.7%, up from 91.2% in Q3.'
    )

    doc.add_heading('Human Resources', level=2)
    doc.add_paragraph(
        'The talent acquisition team filled 127 positions across 8 departments. Average '
        'time-to-fill decreased to 34 days from 41 days in the previous quarter. The '
        'engineering department added 52 software engineers, 12 data scientists, and 8 '
        'DevOps specialists. Employee engagement survey results showed an overall score '
        'of 4.2 out of 5.0, with particularly strong scores in team collaboration (4.5) '
        'and professional development opportunities (4.3).'
    )

    doc.add_page_break()

    # === PAGE 5 ===
    doc.add_heading('Strategic Outlook', level=1)
    doc.add_paragraph(
        'Looking ahead to Q1 2025, the organization is well-positioned to capitalize on '
        'several emerging opportunities. The product roadmap includes the launch of our '
        'AI-powered automation suite, which has shown promising results in beta testing '
        'with 15 pilot customers reporting an average 34% reduction in manual workflow tasks.'
    )
    doc.add_paragraph(
        'Geographic expansion plans include establishing a presence in the Brazilian market '
        'through a partnership with Grupo Votorantim, with an expected launch date of '
        'April 2025. The Latin American market represents a $4.2 billion addressable '
        'opportunity that remains largely untapped by our primary competitors.'
    )
    doc.add_paragraph(
        'Investment priorities for the coming quarter include $8.5 million for cloud '
        'infrastructure migration, $5.2 million for the AI research lab, and $3.7 million '
        'for the customer experience transformation program. These investments are expected '
        'to yield measurable returns within 12-18 months.'
    )

    doc.add_heading('Risk Assessment', level=2)
    doc.add_paragraph(
        'Key risks identified for the upcoming period include potential regulatory changes '
        'in the European data privacy landscape, currency fluctuation impacts on international '
        'revenue (estimated $2.1M exposure), and competitive pressure from three well-funded '
        'startups entering our core market segment. Mitigation strategies have been developed '
        'for each identified risk and are detailed in Appendix C of the full strategic plan.'
    )

    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'Q4 2024 was a transformative quarter for the organization, marked by strong financial '
        'performance, successful product launches, and meaningful progress on our strategic '
        'initiatives. The leadership team remains confident in our ability to sustain growth '
        'momentum while maintaining operational discipline and investing in long-term capabilities.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
