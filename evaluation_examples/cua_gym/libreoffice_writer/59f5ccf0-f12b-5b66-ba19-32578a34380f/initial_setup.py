"""
Initial Setup: Create a Writer document with a formatted warning box paragraph
Task ID: writer_tech_059
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_059'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    heading = doc.add_heading('Technical Documentation Standards', level=1)

    # --- Introductory paragraphs ---
    doc.add_paragraph(
        'This document outlines the standard formatting templates used across '
        'all technical documentation produced by the Engineering Documentation '
        'team at Meridian Systems Inc. All authors must follow these guidelines '
        'when creating user manuals, API references, and internal procedure documents.'
    )

    doc.add_paragraph(
        'Last updated: March 2026. Document owner: Lisa Yamamoto, '
        'Senior Technical Writer.'
    )

    # --- Section: Callout Box Standards ---
    doc.add_heading('Callout Box Standards', level=2)

    doc.add_paragraph(
        'Technical documents frequently use callout boxes to draw attention to '
        'important information. The following paragraph demonstrates the standard '
        'warning box format that should be used consistently across all documents.'
    )

    # --- The Warning Box Paragraph ---
    # This is the key paragraph with: WARNING: prefix (bold red),
    # yellow background shading, and red left border
    warning_para = doc.add_paragraph()

    # Set paragraph-level shading (yellow background) and red left border
    pPr = warning_para._element.get_or_add_pPr()

    # Add yellow background shading
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFFF00"/>')
    pPr.append(shd)

    # Add red left border with padding
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="4" w:color="FF0000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Add indentation for visual offset
    warning_para.paragraph_format.left_indent = Inches(0.25)
    warning_para.paragraph_format.space_before = Pt(6)
    warning_para.paragraph_format.space_after = Pt(6)

    # "WARNING:" prefix - bold, red
    run_prefix = warning_para.add_run('WARNING: ')
    run_prefix.bold = True
    run_prefix.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run_prefix.font.size = Pt(11)
    run_prefix.font.name = 'Liberation Sans'

    # Warning text - normal
    run_text = warning_para.add_run(
        'Improper handling of high-voltage components may result in equipment '
        'damage or personal injury. Always disconnect the power supply and wait '
        'at least 30 seconds before servicing internal modules.'
    )
    run_text.font.size = Pt(11)
    run_text.font.name = 'Liberation Sans'

    # --- More document content ---
    doc.add_heading('Usage Instructions', level=2)

    doc.add_paragraph(
        'To maintain consistency, this warning box format should be saved as an '
        'AutoText entry in LibreOffice Writer. This allows any team member to '
        'quickly insert a properly formatted warning box into their documents '
        'without manually recreating the formatting each time.'
    )

    doc.add_paragraph(
        'The AutoText entry should preserve all formatting attributes including '
        'the red left border, yellow background shading, and the bold red '
        '"WARNING:" prefix text.'
    )

    # --- Section: Other Callout Types ---
    doc.add_heading('Other Standard Callout Types', level=2)

    doc.add_paragraph(
        'In addition to warnings, the following callout types are used:'
    )

    # Note callout (different style for contrast)
    note_para = doc.add_paragraph()
    note_pPr = note_para._element.get_or_add_pPr()
    note_shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="D9E2F3"/>')
    note_pPr.append(note_shd)
    note_bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="4" w:color="4472C4"/>'
        f'</w:pBdr>'
    )
    note_pPr.append(note_bdr)
    note_para.paragraph_format.left_indent = Inches(0.25)
    note_para.paragraph_format.space_before = Pt(6)
    note_para.paragraph_format.space_after = Pt(6)

    run_note_prefix = note_para.add_run('NOTE: ')
    run_note_prefix.bold = True
    run_note_prefix.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run_note_prefix.font.size = Pt(11)
    run_note_prefix.font.name = 'Liberation Sans'

    run_note_text = note_para.add_run(
        'Configuration changes take effect after the next system restart. '
        'Current sessions will not be affected until the service is recycled.'
    )
    run_note_text.font.size = Pt(11)
    run_note_text.font.name = 'Liberation Sans'

    # --- Closing section ---
    doc.add_heading('Document Revision History', level=2)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    headers = ['Version', 'Date', 'Author', 'Changes']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    revisions = [
        ['1.0', '2025-09-15', 'Lisa Yamamoto', 'Initial release'],
        ['1.1', '2025-12-03', 'David Okafor', 'Added callout standards section'],
        ['1.2', '2026-03-10', 'Lisa Yamamoto', 'Updated warning box format'],
    ]
    for r, row_data in enumerate(revisions, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
