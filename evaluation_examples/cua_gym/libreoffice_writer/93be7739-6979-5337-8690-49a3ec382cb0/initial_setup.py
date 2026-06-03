"""
Initial Setup: manuscript_review.docx with three comments for navigation and deletion task
Task ID: writer_struct_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'manuscript_review'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_comment_to_paragraph(doc, para, comment_text, author="Reviewer", date="2025-01-15T10:00:00Z", comment_id=None):
    """Add a comment to a paragraph in the document."""
    from docx.opc.part import XmlPart
    from docx.opc.packuri import PackURI
    from docx.oxml import parse_xml

    part = doc.part
    COMMENTS_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'

    comments_part_obj = None
    for rel in part.rels.values():
        if rel.reltype == COMMENTS_REL_TYPE:
            comments_part_obj = rel.target_part
            break

    if comments_part_obj is None:
        comments_xml = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
</w:comments>'''
        comments_elem = parse_xml(comments_xml)
        comments_part_obj = XmlPart(
            PackURI('/word/comments.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
            comments_elem,
            part.package
        )
        part.relate_to(comments_part_obj, COMMENTS_REL_TYPE)

    # Access the lxml element via XmlPart.element
    comments_root = comments_part_obj.element

    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Create comment element
    comment_elem = etree.SubElement(comments_root, f'{{{ns_w}}}comment')
    comment_elem.set(f'{{{ns_w}}}id', str(comment_id))
    comment_elem.set(f'{{{ns_w}}}author', author)
    comment_elem.set(f'{{{ns_w}}}date', date)
    comment_elem.set(f'{{{ns_w}}}initials', author[:2].upper())

    # Add paragraph with text inside comment
    comment_para = etree.SubElement(comment_elem, f'{{{ns_w}}}p')
    comment_run = etree.SubElement(comment_para, f'{{{ns_w}}}r')
    comment_text_elem = etree.SubElement(comment_run, f'{{{ns_w}}}t')
    comment_text_elem.text = comment_text
    comment_text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Now add comment reference in the paragraph
    # We wrap the paragraph text with commentRangeStart and commentRangeEnd
    para_elem = para._element

    # Insert commentRangeStart before the first run
    comment_range_start = OxmlElement('w:commentRangeStart')
    comment_range_start.set(qn('w:id'), str(comment_id))

    comment_range_end = OxmlElement('w:commentRangeEnd')
    comment_range_end.set(qn('w:id'), str(comment_id))

    # Insert before first run element
    runs = para_elem.findall(f'{{{ns_w}}}r')
    if runs:
        para_elem.insert(list(para_elem).index(runs[0]), comment_range_start)
    else:
        para_elem.append(comment_range_start)

    # Append commentRangeEnd after last run
    para_elem.append(comment_range_end)

    # Add run with commentReference
    ref_run = OxmlElement('w:r')
    ref_rpr = OxmlElement('w:rPr')
    ref_run.append(ref_rpr)
    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), str(comment_id))
    ref_run.append(comment_ref)
    para_elem.append(ref_run)


def create_initial():
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ========== PAGE 1 ==========
    # Title
    title_para = doc.add_heading("Adaptive Learning Strategies in Modern Education: A Comprehensive Review", level=1)

    # Abstract heading
    doc.add_heading("Abstract", level=2)

    # Abstract paragraph - comment 1 anchored here ("Good introduction")
    abstract_para = doc.add_paragraph(
        "This manuscript provides a comprehensive review of adaptive learning strategies in contemporary "
        "educational environments. We examine how technology-mediated instruction, personalized feedback "
        "systems, and evidence-based pedagogical frameworks contribute to improved learning outcomes. "
        "Drawing on data from 47 longitudinal studies conducted between 2018 and 2024, we identify "
        "key patterns that distinguish effective adaptive learning implementations from less successful "
        "approaches. Our findings suggest that successful programs share three core characteristics: "
        "robust formative assessment practices, flexible pacing mechanisms, and strong instructor-student "
        "communication channels."
    )

    doc.add_paragraph(
        "Keywords: adaptive learning, personalized education, formative assessment, technology-enhanced "
        "instruction, educational outcomes"
    )

    # Introduction section
    intro_heading = doc.add_heading("1. Introduction", level=2)

    intro_para1 = doc.add_paragraph(
        "The landscape of modern education has undergone profound transformation over the past two decades. "
        "Driven by advances in educational technology, cognitive science, and data analytics, educators "
        "and researchers have developed increasingly sophisticated approaches to personalized learning. "
        "Central to these developments is the concept of adaptive instruction—teaching methods that "
        "dynamically adjust to the needs, abilities, and progress of individual learners."
    )

    intro_para2 = doc.add_paragraph(
        "Traditional classroom instruction has historically relied on a one-size-fits-all model, "
        "where all students receive identical content, pacing, and assessment regardless of their "
        "individual starting points or learning trajectories. While this approach offers administrative "
        "simplicity, research consistently demonstrates its limitations, particularly for students "
        "at the extremes of the ability spectrum."
    )

    # Page break to go to page 2
    doc.add_page_break()

    # ========== PAGE 2 ==========
    doc.add_heading("2. Literature Review", level=2)

    lit_para1 = doc.add_paragraph(
        "The theoretical foundations of adaptive learning trace back to Bloom's (1984) seminal work "
        "on mastery learning, which demonstrated that virtually all students could achieve high levels "
        "of competence when given adequate time and appropriate instructional support. This foundational "
        "insight has been expanded and refined through decades of subsequent research."
    )

    # Paragraph where comment 2 will be anchored ("Revise this paragraph")
    lit_para2 = doc.add_paragraph(
        "Recent empirical investigations have explored the mechanisms through which adaptive systems "
        "produce their effects. Some researchers emphasize the role of immediate feedback in accelerating "
        "skill acquisition (Johnson & Williams, 2021), while others point to the motivational benefits "
        "of appropriately challenging content (Chen et al., 2022). The integration of artificial "
        "intelligence into adaptive platforms has opened new possibilities for real-time learning "
        "analytics and predictive modeling of student performance trajectories."
    )

    lit_para3 = doc.add_paragraph(
        "However, the evidence base for adaptive learning is not uniformly positive. Several meta-analyses "
        "have identified significant methodological limitations in the existing research, including "
        "small sample sizes, short intervention periods, and a tendency to focus on easily measured "
        "cognitive outcomes at the expense of deeper learning processes and transfer."
    )

    doc.add_heading("2.1 Technological Infrastructure", level=3)

    tech_para = doc.add_paragraph(
        "The implementation of adaptive learning systems requires substantial technological infrastructure. "
        "Learning Management Systems (LMS) must be capable of tracking granular student interactions, "
        "storing large volumes of performance data, and executing complex algorithms that translate "
        "data into instructional decisions. The rapid evolution of cloud computing and machine learning "
        "capabilities has dramatically lowered the technical barriers to implementing sophisticated "
        "adaptive systems."
    )

    # Page break to page 3
    doc.add_page_break()

    # ========== PAGE 3 ==========
    doc.add_heading("3. Methodology", level=2)

    method_para1 = doc.add_paragraph(
        "This review employed a systematic search strategy across multiple academic databases including "
        "ERIC, PsycINFO, Web of Science, and Scopus. Search terms included combinations of 'adaptive "
        "learning,' 'personalized instruction,' 'intelligent tutoring systems,' 'formative assessment,' "
        "and 'learning analytics.' The search was limited to peer-reviewed publications from 2010 "
        "to 2024 to ensure relevance to contemporary practice."
    )

    method_para2 = doc.add_paragraph(
        "Studies were included if they reported empirical data on student learning outcomes in K-12 "
        "or higher education settings, utilized a comparison or control condition, and provided "
        "sufficient methodological detail for quality assessment. Studies were excluded if they "
        "relied solely on self-reported data, lacked clear outcome measures, or involved populations "
        "outside educational contexts."
    )

    doc.add_heading("3.1 Quality Assessment", level=3)

    quality_para = doc.add_paragraph(
        "Each included study was evaluated using the Mixed Methods Appraisal Tool (MMAT), "
        "which provides a framework for assessing the methodological quality of quantitative, "
        "qualitative, and mixed-methods research. Two independent reviewers assessed each study, "
        "with disagreements resolved through discussion or arbitration by a third reviewer. "
        "Inter-rater reliability was calculated using Cohen's kappa coefficient."
    )

    doc.add_heading("3.2 Data Synthesis", level=3)

    synthesis_para = doc.add_paragraph(
        "Given the heterogeneity of included studies in terms of populations, interventions, "
        "and outcome measures, we employed a narrative synthesis approach rather than quantitative "
        "meta-analysis. This approach allowed us to integrate findings across diverse methodological "
        "traditions while acknowledging the limitations of direct comparison."
    )

    # Page break to page 4
    doc.add_page_break()

    # ========== PAGE 4 ==========
    doc.add_heading("4. Results", level=2)

    results_para1 = doc.add_paragraph(
        "Our systematic search identified 847 potentially relevant publications. Following title "
        "and abstract screening, 312 full-text articles were reviewed. Of these, 47 studies met "
        "all inclusion criteria and were included in the final synthesis. The included studies "
        "represented a diverse range of educational contexts, including elementary schools (n=12), "
        "secondary schools (n=18), undergraduate programs (n=14), and professional training "
        "environments (n=3)."
    )

    # Paragraph where comment 3 anchored ("Strong evidence presented")
    results_para2 = doc.add_paragraph(
        "Across the 47 included studies, adaptive learning interventions demonstrated consistently "
        "positive effects on student achievement outcomes. Effect sizes ranged from d=0.34 to d=1.12, "
        "with a median effect size of d=0.67, which is considered a medium-to-large effect in "
        "educational research (Cohen, 1988). These findings are consistent with prior reviews "
        "of educational technology interventions and suggest that well-implemented adaptive "
        "systems can produce meaningful improvements in student learning."
    )

    results_para3 = doc.add_paragraph(
        "Moderator analyses revealed that several factors significantly influenced the magnitude "
        "of adaptive learning effects. The most consistent moderator was implementation fidelity: "
        "studies reporting high fidelity to the intended intervention design showed effect sizes "
        "approximately 40% larger than those reporting moderate or low fidelity. This finding "
        "underscores the importance of adequate teacher training and ongoing technical support."
    )

    doc.add_heading("4.1 Qualitative Themes", level=3)

    themes_para = doc.add_paragraph(
        "Analysis of qualitative data from case studies and interview-based research revealed "
        "several recurring themes. Teachers consistently reported that adaptive systems helped "
        "them identify struggling students earlier and more accurately than traditional assessment "
        "methods. Students expressed appreciation for the ability to work at their own pace and "
        "receive immediate feedback on their performance."
    )

    # Page break to page 5
    doc.add_page_break()

    # ========== PAGE 5 ==========
    doc.add_heading("5. Discussion", level=2)

    discuss_para1 = doc.add_paragraph(
        "The results of this review provide strong support for the effectiveness of adaptive "
        "learning interventions across educational contexts. The consistency of positive findings "
        "across diverse populations, settings, and technological platforms suggests that the "
        "benefits of adaptive learning reflect genuine instructional advantages rather than "
        "artifacts of specific implementation contexts."
    )

    discuss_para2 = doc.add_paragraph(
        "At the same time, our analysis reveals important nuances that should inform both research "
        "and practice. The substantial variability in effect sizes suggests that not all adaptive "
        "learning implementations are equally effective. Understanding the conditions that maximize "
        "the benefits of adaptive learning is critical for both researchers seeking to advance "
        "the field and practitioners seeking to improve student outcomes."
    )

    doc.add_heading("5.1 Implications for Practice", level=3)

    practice_para = doc.add_paragraph(
        "Educational leaders considering the adoption of adaptive learning systems should prioritize "
        "comprehensive teacher preparation, including both technical training on the specific platform "
        "and professional development on interpreting and responding to learning analytics data. "
        "Our review suggests that the quality of teacher support is a stronger predictor of outcomes "
        "than the sophistication of the underlying technology."
    )

    doc.add_heading("6. Conclusion", level=2)

    conclusion_para = doc.add_paragraph(
        "This comprehensive review of 47 empirical studies provides compelling evidence for the "
        "effectiveness of adaptive learning systems in improving student achievement. Our analysis "
        "highlights the critical role of implementation quality, teacher support, and contextual "
        "factors in moderating these effects. Future research should focus on identifying the "
        "specific mechanisms through which adaptive systems produce their effects, with particular "
        "attention to motivational and metacognitive pathways."
    )

    references_heading = doc.add_heading("References", level=2)

    doc.add_paragraph(
        "Bloom, B. S. (1984). The 2 sigma problem: The search for methods of group instruction "
        "as effective as one-to-one tutoring. Educational Researcher, 13(6), 4-16."
    )
    doc.add_paragraph(
        "Chen, L., Park, J., & Martinez, R. (2022). Motivational effects of challenge-calibrated "
        "adaptive content in online learning environments. Journal of Educational Psychology, "
        "114(3), 512-528."
    )
    doc.add_paragraph(
        "Johnson, M. & Williams, T. (2021). Immediate feedback mechanisms in intelligent tutoring "
        "systems: A meta-analysis. Review of Educational Research, 91(2), 201-237."
    )

    # Save document temporarily before adding comments
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created (pre-comment): {OUTPUT}')

    # Re-open and add comments via XML manipulation
    doc2 = Document(OUTPUT)

    # Get paragraphs from the document
    all_paras = doc2.paragraphs

    # Find target paragraphs by their text content
    abstract_target = None
    lit_review_target = None
    results_target = None

    for p in all_paras:
        txt = p.text
        if abstract_target is None and 'This manuscript provides a comprehensive review' in txt:
            abstract_target = p
        if lit_review_target is None and 'Recent empirical investigations have explored' in txt:
            lit_review_target = p
        if results_target is None and 'Across the 47 included studies, adaptive learning' in txt:
            results_target = p

    print(f'Found abstract_target: {abstract_target is not None}')
    print(f'Found lit_review_target: {lit_review_target is not None}')
    print(f'Found results_target: {results_target is not None}')

    # Add comments to found paragraphs
    if abstract_target:
        add_comment_to_paragraph(doc2, abstract_target, "Good introduction",
                                  author="Dr. Sarah Mitchell", date="2025-01-15T09:30:00Z", comment_id=1)
    if lit_review_target:
        add_comment_to_paragraph(doc2, lit_review_target, "Revise this paragraph",
                                  author="Dr. Sarah Mitchell", date="2025-01-15T10:15:00Z", comment_id=2)
    if results_target:
        add_comment_to_paragraph(doc2, results_target, "Strong evidence presented",
                                  author="Dr. Sarah Mitchell", date="2025-01-15T11:00:00Z", comment_id=3)

    doc2.save(OUTPUT)
    print(f'Initial file with comments saved: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
