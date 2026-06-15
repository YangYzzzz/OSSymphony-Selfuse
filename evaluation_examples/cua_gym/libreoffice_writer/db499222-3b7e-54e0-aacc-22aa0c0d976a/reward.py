"""
Reward Script: Apply 'Box List Blue' AutoFormat style to comparison table
Task ID: writer_tbl_018
Domain: libreoffice_writer
Scoring:
  Component 1: Header row (row 0) cells have dark blue shading (~#2E74B5) — 0.4 points
  Component 2: Header row text is bold and white (~#FFFFFF) — 0.3 points
  Component 3: Data rows have blue-themed alternating shading (light blue / white) — 0.3 points
Total: 1.0

Ground truth: The 'Box List Blue' AutoFormat table style is applied to the table,
giving it blue-themed formatting with styled header row, alternating colors, and
appropriate borders as defined by the style.

Key changes from initial to golden:
  - Initial: Table has basic 'Table Grid' style with no cell shading, default fonts
  - Golden: Cells have explicit blue shading applied:
      Row 0 (header): fill=#2E74B5 (dark blue), text bold+white
      Row 1: fill=#DEEAF1 (light blue)
      Row 2: fill=#FFFFFF (white)
      Row 3: fill=#DEEAF1 (light blue)
"""

import os
from docx import Document
from docx.oxml.ns import qn
from math import sqrt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_018'
FILE_PATH = f'{WORKDIR}/comparison.docx'


def hex_to_rgb(hex_str):
    """Convert hex color string (e.g. '2E74B5') to (R, G, B) tuple."""
    try:
        h = hex_str.strip().lstrip('#')
        return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except Exception:
        return None


def color_distance(c1, c2):
    """Euclidean distance between two (R, G, B) tuples."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def get_cell_fill(cell):
    """Return fill hex string of a cell, or None if no shading."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    return fill if fill and fill.upper() not in ('AUTO', 'NONE', '') else None


def verify_task(file_path):
    """
    Verify that 'Box List Blue' AutoFormat style was applied to the comparison table.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — precondition gate
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have exactly 1 table with 4 rows and 3 columns
    try:
        if len(doc.tables) == 0:
            print("CRITICAL: No tables found in document")
            print("REWARD: 0.0")
            return 0.0
        table = doc.tables[0]
        if len(table.rows) != 4 or len(table.columns) != 3:
            print(f"CRITICAL: Expected 4x3 table, found {len(table.rows)}x{len(table.columns)}")
            print("REWARD: 0.0")
            return 0.0
    except Exception as e:
        print(f"CRITICAL: Cannot inspect table structure: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Header row (row 0) cells have dark blue shading ~#2E74B5 (0.4 points)
    # This FAILS on initial (no shading) and PASSES on golden (fill=2E74B5)
    HEADER_BLUE = hex_to_rgb('2E74B5')  # dark blue header from Box List Blue style
    TOLERANCE = 50  # RGB Euclidean distance tolerance
    try:
        header_row = table.rows[0]
        header_blue_cells = 0
        header_details = []
        for j, cell in enumerate(header_row.cells):
            fill = get_cell_fill(cell)
            if fill is not None:
                fill_rgb = hex_to_rgb(fill)
                if fill_rgb is not None:
                    dist = color_distance(fill_rgb, HEADER_BLUE)
                    if dist <= TOLERANCE:
                        header_blue_cells += 1
                        header_details.append(f'cell({0},{j}): fill=#{fill} (dist={dist:.1f})')
                    else:
                        header_details.append(f'cell({0},{j}): fill=#{fill} (dist={dist:.1f}, too far from 2E74B5)')
                else:
                    header_details.append(f'cell({0},{j}): fill={fill} (could not parse)')
            else:
                header_details.append(f'cell({0},{j}): no fill')

        if header_blue_cells == 3:
            print(f"PASS: Component 1 — Header row has dark blue shading on all 3 cells: {header_details} (0.4 pts)")
            total_score += 0.4
        elif header_blue_cells > 0:
            print(f"PARTIAL: Component 1 — Only {header_blue_cells}/3 header cells have dark blue shading: {header_details}")
        else:
            print(f"FAIL: Component 1 — Header row has no dark blue (#2E74B5) shading. Details: {header_details}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row text is bold and white (~#FFFFFF) (0.3 points)
    # This FAILS on initial (bold=None, color=None) and PASSES on golden (bold=True, color=FFFFFF)
    WHITE_RGB = (255, 255, 255)
    try:
        header_row = table.rows[0]
        white_bold_cells = 0
        font_details = []
        for j, cell in enumerate(header_row.cells):
            cell_bold = False
            cell_white = False
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        if run.bold is True:
                            cell_bold = True
                        try:
                            if run.font.color and run.font.color.rgb is not None:
                                rgb = run.font.color.rgb
                                # RGBColor supports indexing [0], [1], [2] or str access
                                # Convert to int via hex string
                                hex_str = str(rgb)  # e.g. 'FFFFFF'
                                color_tuple = hex_to_rgb(hex_str)
                                if color_tuple is not None and color_distance(color_tuple, WHITE_RGB) <= TOLERANCE:
                                    cell_white = True
                        except Exception:
                            pass
            if cell_bold and cell_white:
                white_bold_cells += 1
                font_details.append(f'cell(0,{j}): bold+white OK')
            else:
                font_details.append(f'cell(0,{j}): bold={cell_bold}, white={cell_white}')

        if white_bold_cells == 3:
            print(f"PASS: Component 2 — Header row text is bold and white on all 3 cells: {font_details} (0.3 pts)")
            total_score += 0.3
        elif white_bold_cells > 0:
            print(f"PARTIAL: Component 2 — Only {white_bold_cells}/3 header cells have bold+white text: {font_details}")
        else:
            print(f"FAIL: Component 2 — Header row text is not bold+white. Details: {font_details}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Data rows have blue-themed alternating shading (light blue / white) (0.3 points)
    # Row 1: fill=#DEEAF1 (light blue), Row 2: fill=#FFFFFF (white), Row 3: fill=#DEEAF1 (light blue)
    # This FAILS on initial (no shading) and PASSES on golden
    LIGHT_BLUE = hex_to_rgb('DEEAF1')  # light blue alternating rows
    WHITE_HEX = hex_to_rgb('FFFFFF')
    try:
        # Expected fills for rows 1, 2, 3
        expected_fills = {
            1: ('DEEAF1', LIGHT_BLUE),
            2: ('FFFFFF', WHITE_HEX),
            3: ('DEEAF1', LIGHT_BLUE),
        }
        rows_pass = 0
        alt_details = []
        for row_idx, (expected_hex, expected_rgb) in expected_fills.items():
            row = table.rows[row_idx]
            row_cells_pass = 0
            for j, cell in enumerate(row.cells):
                fill = get_cell_fill(cell)
                if fill is not None:
                    fill_rgb = hex_to_rgb(fill)
                    if fill_rgb is not None:
                        dist = color_distance(fill_rgb, expected_rgb)
                        if dist <= TOLERANCE:
                            row_cells_pass += 1
                        else:
                            alt_details.append(f'row{row_idx} cell{j}: fill=#{fill} (expected #{expected_hex}, dist={dist:.1f})')
                    else:
                        alt_details.append(f'row{row_idx} cell{j}: fill={fill} (parse error)')
                else:
                    alt_details.append(f'row{row_idx} cell{j}: no fill (expected #{expected_hex})')
            if row_cells_pass == 3:
                rows_pass += 1
                alt_details.append(f'row{row_idx}: all 3 cells #{expected_hex} OK')

        if rows_pass == 3:
            print(f"PASS: Component 3 — Data rows have correct alternating blue shading: {alt_details} (0.3 pts)")
            total_score += 0.3
        elif rows_pass > 0:
            print(f"PARTIAL: Component 3 — Only {rows_pass}/3 data rows have correct alternating shading: {alt_details}")
        else:
            print(f"FAIL: Component 3 — Data rows do not have expected alternating blue shading. Details: {alt_details}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
