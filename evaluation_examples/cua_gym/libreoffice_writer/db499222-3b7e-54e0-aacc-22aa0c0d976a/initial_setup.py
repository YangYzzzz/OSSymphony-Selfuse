"""
Initial Setup: Create comparison.docx with a plain table (no autoformat styling)
Task ID: writer_tbl_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'writer_tbl_018'
OUTPUT = f'{DESKTOP}/comparison.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Add a brief intro paragraph
    intro = doc.add_paragraph("Below is a comparison of our available service plans.")

    # Create table: 4 rows x 3 columns with default "Table Grid" style
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Row 1 (header row)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Feature'
    header_cells[1].text = 'Plan A'
    header_cells[2].text = 'Plan B'

    # Row 2
    row2 = table.rows[1].cells
    row2[0].text = 'Storage'
    row2[1].text = '10 GB'
    row2[2].text = '50 GB'

    # Row 3
    row3 = table.rows[2].cells
    row3[0].text = 'Users'
    row3[1].text = '5'
    row3[2].text = '25'

    # Row 4
    row4 = table.rows[3].cells
    row4[0].text = 'Price'
    row4[1].text = '$10/mo'
    row4[2].text = '$30/mo'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
