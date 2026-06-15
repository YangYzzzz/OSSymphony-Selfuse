"""
Initial Setup: Create onboarding form document with 'Department: ' text
Task ID: writer_tm_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_058'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('New Employee Onboarding Form', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.italic = True

    # --- Horizontal line ---
    doc.add_paragraph('_' * 60)

    # --- Section 1: Personal Information ---
    doc.add_heading('Section 1: Personal Information', level=1)

    p1 = doc.add_paragraph()
    r1 = p1.add_run('Full Name: ')
    r1.bold = True
    r1.font.size = Pt(11)
    p1.add_run('________________________________')

    p2 = doc.add_paragraph()
    r2 = p2.add_run('Employee ID: ')
    r2.bold = True
    r2.font.size = Pt(11)
    p2.add_run('________________________________')

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Email Address: ')
    r3.bold = True
    r3.font.size = Pt(11)
    p3.add_run('________________________________')

    p4 = doc.add_paragraph()
    r4 = p4.add_run('Phone Number: ')
    r4.bold = True
    r4.font.size = Pt(11)
    p4.add_run('________________________________')

    # --- Section 2: Position Details ---
    doc.add_heading('Section 2: Position Details', level=1)

    p5 = doc.add_paragraph()
    r5 = p5.add_run('Job Title: ')
    r5.bold = True
    r5.font.size = Pt(11)
    p5.add_run('________________________________')

    # Department line - NO input field, just the label with blank space
    p6 = doc.add_paragraph()
    r6 = p6.add_run('Department: ')
    r6.bold = True
    r6.font.size = Pt(11)

    p7 = doc.add_paragraph()
    r7 = p7.add_run('Start Date: ')
    r7.bold = True
    r7.font.size = Pt(11)
    p7.add_run('________________________________')

    p8 = doc.add_paragraph()
    r8 = p8.add_run('Reporting Manager: ')
    r8.bold = True
    r8.font.size = Pt(11)
    p8.add_run('________________________________')

    # --- Section 3: Emergency Contact ---
    doc.add_heading('Section 3: Emergency Contact', level=1)

    p9 = doc.add_paragraph()
    r9 = p9.add_run('Contact Name: ')
    r9.bold = True
    r9.font.size = Pt(11)
    p9.add_run('________________________________')

    p10 = doc.add_paragraph()
    r10 = p10.add_run('Relationship: ')
    r10.bold = True
    r10.font.size = Pt(11)
    p10.add_run('________________________________')

    p11 = doc.add_paragraph()
    r11 = p11.add_run('Contact Phone: ')
    r11.bold = True
    r11.font.size = Pt(11)
    p11.add_run('________________________________')

    # --- Section 4: Acknowledgment ---
    doc.add_heading('Section 4: Acknowledgment', level=1)

    ack = doc.add_paragraph(
        'I acknowledge that I have received and reviewed the employee handbook, '
        'company policies, and safety guidelines. I understand that it is my '
        'responsibility to comply with all organizational rules and procedures.'
    )
    ack.paragraph_format.space_after = Pt(12)

    sig = doc.add_paragraph()
    r_sig = sig.add_run('Signature: ')
    r_sig.bold = True
    r_sig.font.size = Pt(11)
    sig.add_run('________________________________')

    date_line = doc.add_paragraph()
    r_date = date_line.add_run('Date: ')
    r_date.bold = True
    r_date.font.size = Pt(11)
    date_line.add_run('________________________________')

    # --- Footer note ---
    doc.add_paragraph('')
    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_fn = footer_note.add_run('Please submit this form to HR within 3 business days of your start date.')
    r_fn.font.size = Pt(9)
    r_fn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r_fn.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
