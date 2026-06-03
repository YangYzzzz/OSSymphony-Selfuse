"""
Reward Script: Mail merge form letter for rental agreement renewals
Task ID: writer_mt_040
Domain: libreoffice_writer
Scoring:
  Component 1: Document has substantial letter content (0.10)
  Component 2: All 5 required mail merge fields present (0.50)
  Component 3: Letter addresses TenantName (Dear/addressed) (0.10)
  Component 4: Letter references 'Unit <UnitNumber>' (0.10)
  Component 5: Monthly rent shown as '$<MonthlyRent>' (0.10)
  Component 6: 'Property Management Office' signature line (0.10)
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_040'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def find_merge_fields_in_xml(doc):
    """Extract all MERGEFIELD names from the document XML."""
    body_xml = doc.element.body.xml
    fields = re.findall(r'MERGEFIELD\s+(\w+)', body_xml)
    return fields


def get_full_text(doc):
    """Get all paragraph text concatenated."""
    return '\n'.join(p.text for p in doc.paragraphs)


def verify_task(file_path):
    """
    Verify mail merge form letter completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    full_text = get_full_text(doc)
    merge_fields = find_merge_fields_in_xml(doc)
    unique_fields = set(merge_fields)

    # Component 1: Document has substantial letter content (0.10 points)
    # Initial doc is blank (1 empty paragraph). Golden has 30 paragraphs with letter content.
    # We check that the document has at least 10 non-empty paragraphs (clearly a letter, not blank).
    try:
        non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
        num_non_empty = len(non_empty_paras)
        if num_non_empty >= 10:
            print(f"PASS: Component 1 - Document has {num_non_empty} non-empty paragraphs (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 - Expected >= 10 non-empty paragraphs, found {num_non_empty}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: All 5 required mail merge fields present in document XML (0.50 points)
    # Each field is worth 0.10 points. Fields must be actual MERGEFIELD entries in XML,
    # not just guillemet text.
    required_fields = ['TenantName', 'UnitNumber', 'MonthlyRent', 'LeaseStartDate', 'LeaseEndDate']
    try:
        fields_found = 0
        for field_name in required_fields:
            if field_name in unique_fields:
                print(f"PASS: Component 2.{fields_found+1} - MERGEFIELD '{field_name}' found in XML (0.10 pts)")
                fields_found += 1
                total_score += 0.10
            else:
                print(f"FAIL: Component 2.{fields_found+1} - MERGEFIELD '{field_name}' NOT found in XML")
        print(f"  Component 2 subtotal: {fields_found}/5 fields found ({fields_found * 0.10:.2f} pts)")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Letter addresses TenantName - "Dear <TenantName>" pattern (0.10 points)
    # The golden document has "Dear <<TenantName>>," as a greeting line.
    # We verify there is a paragraph containing "Dear" AND the TenantName merge field.
    try:
        greeting_count = sum(
            1 for p in doc.paragraphs
            if 'dear' in p.text.lower() and 'MERGEFIELD TenantName' in p._element.xml
        )
        if greeting_count > 0:
            print(f"PASS: Component 3 - Letter addresses TenantName with 'Dear' greeting (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 - No paragraph with 'Dear' + MERGEFIELD TenantName found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Letter references 'Unit <UnitNumber>' (0.10 points)
    # The golden document mentions "Unit <<UnitNumber>>" in multiple places.
    # We check for a paragraph containing "Unit" followed by the UnitNumber merge field.
    try:
        unit_ref_count = sum(
            1 for p in doc.paragraphs
            if 'unit' in p.text.lower() and 'MERGEFIELD UnitNumber' in p._element.xml
        )
        if unit_ref_count > 0:
            print(f"PASS: Component 4 - Letter references 'Unit <UnitNumber>' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 - No paragraph with 'Unit' + MERGEFIELD UnitNumber found")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Monthly rent shown as '$<MonthlyRent>' (0.10 points)
    # The golden document has "Monthly Rent: $<<MonthlyRent>>"
    # We check for a paragraph containing '$' and MERGEFIELD MonthlyRent.
    try:
        rent_ref_count = sum(
            1 for p in doc.paragraphs
            if '$' in p.text and 'MERGEFIELD MonthlyRent' in p._element.xml
        )
        if rent_ref_count > 0:
            print(f"PASS: Component 5 - Monthly rent shown as '$<MonthlyRent>' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 - No paragraph with '$' + MERGEFIELD MonthlyRent found")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: 'Property Management Office' signature line (0.10 points)
    # The golden document has "Property Management Office" near the end as a signature.
    try:
        sig_count = sum(
            1 for p in doc.paragraphs
            if 'property management office' in p.text.lower()
        )
        if sig_count > 0:
            print(f"PASS: Component 6 - 'Property Management Office' signature found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 - 'Property Management Office' signature not found")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
