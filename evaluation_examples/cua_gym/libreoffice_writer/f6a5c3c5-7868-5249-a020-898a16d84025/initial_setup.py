"""
Initial Setup: Create a Writer document with 10 bulleted items using default 'List Bullet' style.
Task ID: writer_bs_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    doc.add_heading('Q1 2025 Action Items', level=1)

    # 10 realistic bulleted items using List Bullet style
    bullet_items = [
        'Review quarterly sales report and identify top-performing regions',
        'Schedule one-on-one meetings with all direct reports by March 15',
        'Submit updated budget proposal to the finance department',
        'Complete onboarding documentation for new team members in Engineering',
        'Coordinate with marketing team on product launch timeline for April',
        'Update the employee handbook with revised remote work policies',
        'Prepare presentation slides for the board meeting on March 28',
        'Audit vendor contracts and flag any renewals due before June 30',
        'Organize team-building event for the Portland office relocation',
        'Finalize performance review criteria with HR before end of quarter',
    ]

    for item in bullet_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
