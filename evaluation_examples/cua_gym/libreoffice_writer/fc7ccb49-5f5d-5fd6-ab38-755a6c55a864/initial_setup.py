"""
Initial Setup: Campaign materials document with page-number-only footer (no disclaimer)
Task ID: writer_mktg_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'campaign_materials'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section):
    """Add a footer with only a centered page number field."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Page number field
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = ' PAGE '
    r2._element.append(instr)

    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Page setup: US Letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---- Page 1: Executive Summary / Campaign Overview ----
    h = doc.add_heading('Apex Dynamics, Inc. — Q2 Campaign Creative Brief', level=1)
    doc.add_paragraph(
        'Prepared by: Marketing Strategy Division\n'
        'Date: March 1, 2025\n'
        'Version: 2.4 (Internal Review)'
    )
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'This document outlines the creative direction, media strategy, and budget allocation '
        'for the Apex Dynamics Q2 2025 product launch campaign. The campaign will target enterprise '
        'software buyers across North America and EMEA, with a focus on the new CloudSync Pro platform.'
    )
    doc.add_paragraph(
        'Key objectives:\n'
        '• Increase brand awareness in the mid-market enterprise segment by 25%\n'
        '• Generate 4,500 qualified leads through digital and event channels\n'
        '• Achieve a 3.2x return on ad spend (ROAS) across paid media channels\n'
        '• Drive 18,000 free trial sign-ups for CloudSync Pro'
    )
    doc.add_page_break()

    # ---- Page 2: Campaign Messaging & Creative Direction ----
    doc.add_heading('Campaign Messaging & Creative Direction', level=1)
    doc.add_heading('Primary Value Proposition', level=2)
    doc.add_paragraph(
        '"CloudSync Pro: Enterprise collaboration without boundaries." '
        'The creative team has developed three distinct message pillars that will anchor all '
        'campaign materials across channels:'
    )
    doc.add_paragraph('Pillar 1: Speed — Deploy in under 48 hours with zero downtime', style='List Number')
    doc.add_paragraph('Pillar 2: Security — SOC 2 Type II certified with end-to-end encryption', style='List Number')
    doc.add_paragraph('Pillar 3: Scale — Supports teams of 10 to 100,000 users seamlessly', style='List Number')

    doc.add_heading('Creative Assets Required', level=2)
    doc.add_paragraph(
        'The following creative assets have been commissioned from the internal design team '
        'and agency partner Meridian Creative Group:'
    )
    data_creative = [
        ['Asset Type', 'Format', 'Quantity', 'Deadline', 'Owner'],
        ['Hero Banner', '1920×1080 px', '3 variants', 'March 15, 2025', 'Internal Design'],
        ['Social Cards', '1080×1080 px', '12 variants', 'March 12, 2025', 'Meridian Creative'],
        ['Video Spot (30s)', 'MP4 H.264', '2 variants', 'March 20, 2025', 'Meridian Creative'],
        ['Email Header', '600×200 px', '4 variants', 'March 10, 2025', 'Internal Design'],
        ['Display Ads', 'Multiple IAB', '8 sizes', 'March 18, 2025', 'Meridian Creative'],
    ]
    table1 = doc.add_table(rows=len(data_creative), cols=5)
    table1.style = 'Table Grid'
    for i, row_data in enumerate(data_creative):
        for j, val in enumerate(row_data):
            cell = table1.cell(i, j)
            run = cell.paragraphs[0].add_run(val)
            if i == 0:
                run.bold = True
    doc.add_page_break()

    # ---- Page 3: Target Audience & Segmentation ----
    doc.add_heading('Target Audience & Segmentation', level=1)
    doc.add_heading('Primary Audience', level=2)
    doc.add_paragraph(
        'The primary audience for the CloudSync Pro campaign consists of IT decision-makers '
        'and C-suite executives at companies with 200–5,000 employees. Firmographic targeting '
        'will prioritize technology, financial services, and healthcare verticals.'
    )
    doc.add_heading('Audience Personas', level=2)

    personas = [
        ('Persona A: The Pragmatic CTO', 'Elena Vasquez, 44, CTO at a 600-person fintech firm. '
         'Prioritizes security, compliance, and integration with existing DevOps toolchains. '
         'Reads TechCrunch, attends AWS re:Invent, and relies heavily on peer recommendations.'),
        ('Persona B: The Growth-Focused VP of IT', 'Derek Thornton, 38, VP IT at a healthcare SaaS company. '
         'Responsible for scaling infrastructure as the company doubles headcount. '
         'Values uptime SLAs, 24/7 support, and transparent pricing.'),
        ('Persona C: The Digital-Native COO', 'Priya Nair, 35, COO at a Series C logistics startup. '
         'Drives cross-functional collaboration and views technology as a competitive differentiator. '
         'Highly active on LinkedIn and participates in startup accelerator networks.'),
    ]
    for title, description in personas:
        doc.add_heading(title, level=3)
        doc.add_paragraph(description)
    doc.add_page_break()

    # ---- Page 4: Digital Media Plan ----
    doc.add_heading('Digital Media Plan', level=1)
    doc.add_paragraph(
        'The digital media plan covers paid search, programmatic display, social advertising, '
        'and content syndication for the period April 1 – June 30, 2025.'
    )
    doc.add_heading('Paid Search (Google & Bing)', level=2)
    doc.add_paragraph(
        'Campaigns will target high-intent keywords across three tiers: '
        'branded, competitor, and category. Expected CPC ranges from $12 to $38 for enterprise '
        'software terms. Total allocation: $185,000.'
    )
    doc.add_heading('Social Advertising', level=2)
    doc.add_paragraph(
        'LinkedIn will serve as the primary social channel, targeting by job function, '
        'company size, and industry. Campaign formats include Sponsored Content, InMail, '
        'and Dynamic Ads. Secondary social spend on Meta will focus on retargeting audiences '
        'identified via the Apex Dynamics website pixel. Total allocation: $210,000.'
    )
    doc.add_heading('Programmatic Display', level=2)
    doc.add_paragraph(
        'Demand-side platform (DSP) buys via The Trade Desk will cover brand awareness '
        'and retargeting objectives. CPM targets: $18–$24 for prospecting, $8–$12 for retargeting. '
        'Total allocation: $95,000.'
    )
    doc.add_page_break()

    # ---- Page 5: Event & Field Marketing ----
    doc.add_heading('Event & Field Marketing Strategy', level=1)
    doc.add_paragraph(
        'In-person and virtual events represent a critical channel for pipeline generation. '
        'The following events have been confirmed for Q2 2025:'
    )
    event_data = [
        ['Event', 'Location', 'Date', 'Expected Attendance', 'Budget'],
        ['Gartner IT Symposium', 'Las Vegas, NV', 'April 7–10, 2025', '12,000+', '$85,000'],
        ['CloudWorld Summit', 'San Francisco, CA', 'April 22–24, 2025', '6,500+', '$62,000'],
        ['Enterprise Connect', 'Orlando, FL', 'May 5–8, 2025', '9,000+', '$74,000'],
        ['Apex Dynamics User Conference', 'Austin, TX', 'May 20–21, 2025', '2,200', '$120,000'],
        ['Infosecurity Europe', 'London, UK', 'June 3–5, 2025', '19,000+', '$55,000'],
        ['SaaStr Annual', 'San Mateo, CA', 'June 10–12, 2025', '10,000+', '$48,000'],
    ]
    table2 = doc.add_table(rows=len(event_data), cols=5)
    table2.style = 'Table Grid'
    for i, row_data in enumerate(event_data):
        for j, val in enumerate(row_data):
            cell = table2.cell(i, j)
            run = cell.paragraphs[0].add_run(val)
            if i == 0:
                run.bold = True
    doc.add_page_break()

    # ---- Page 6: Content Marketing & SEO ----
    doc.add_heading('Content Marketing & SEO Strategy', level=1)
    doc.add_heading('Content Pillars', level=2)
    doc.add_paragraph(
        'The content team will produce a steady cadence of thought leadership, product education, '
        'and customer success content throughout Q2. All content maps to buyer journey stages: '
        'Awareness, Consideration, and Decision.'
    )
    content_types = [
        ('Whitepapers (3)', 'Deep-dive technical guides on enterprise integration patterns, '
         'zero-trust security architecture, and ROI modeling for cloud collaboration tools.'),
        ('Case Studies (6)', 'Customer success stories from Meridian Financial, NovaTech Solutions, '
         'Brightside Healthcare, and three additional enterprise accounts.'),
        ('Webinars (4)', 'Monthly live sessions: "CloudSync Pro Product Demo", "Security Best Practices", '
         '"Migration Guide from Legacy Tools", "Customer Panel: Real-World Results".'),
        ('Blog Posts (24)', 'Two posts per week across topics including IT trends, product updates, '
         'industry analysis, and how-to guides. Target 3,000–5,000 words for pillar content.'),
    ]
    for ct_title, ct_desc in content_types:
        p = doc.add_paragraph(style='List Bullet')
        run_title = p.add_run(ct_title + ': ')
        run_title.bold = True
        p.add_run(ct_desc)
    doc.add_page_break()

    # ---- Page 7: Campaign Budget Summary ----
    doc.add_heading('Campaign Budget Summary', level=1)
    doc.add_paragraph(
        'The following table summarizes the total Q2 2025 campaign budget by channel. '
        'All figures are in USD. Final approval pending CFO sign-off by March 14, 2025.'
    )
    budget_data = [
        ['Channel', 'Allocated Budget', '% of Total', 'Lead Goal', 'Cost per Lead Target'],
        ['Paid Search (SEM)', '$185,000', '17.5%', '850', '$218'],
        ['Social Advertising (LinkedIn/Meta)', '$210,000', '19.9%', '1,200', '$175'],
        ['Programmatic Display', '$95,000', '9.0%', '400', '$238'],
        ['Event & Field Marketing', '$444,000', '42.0%', '1,600', '$278'],
        ['Content Marketing & SEO', '$68,000', '6.4%', '300', '$227'],
        ['Email Marketing', '$24,000', '2.3%', '150', '$160'],
        ['Influencer & Partner Co-Marketing', '$31,000', '2.9%', 'N/A', 'N/A'],
        ['TOTAL', '$1,057,000', '100%', '4,500', '$235'],
    ]
    table3 = doc.add_table(rows=len(budget_data), cols=5)
    table3.style = 'Table Grid'
    for i, row_data in enumerate(budget_data):
        for j, val in enumerate(row_data):
            cell = table3.cell(i, j)
            run = cell.paragraphs[0].add_run(val)
            if i == 0 or (i == len(budget_data) - 1):
                run.bold = True
    doc.add_page_break()

    # ---- Page 8: KPIs, Reporting & Approvals ----
    doc.add_heading('KPIs, Reporting Schedule & Approvals', level=1)
    doc.add_heading('Key Performance Indicators', level=2)
    kpi_data = [
        ['KPI', 'Target', 'Measurement Tool', 'Reporting Frequency'],
        ['Total Qualified Leads', '4,500', 'Salesforce CRM', 'Weekly'],
        ['Free Trial Sign-Ups', '18,000', 'Product Analytics (Mixpanel)', 'Daily'],
        ['Paid Media ROAS', '3.2x', 'Google Analytics 4 / LinkedIn Insight', 'Weekly'],
        ['Brand Awareness Lift', '+25%', 'Nielsen Brand Survey (post-campaign)', 'End of Q2'],
        ['Email Open Rate', '28%+', 'Marketo', 'Per Campaign'],
        ['Content-Sourced Pipeline', '$3.2M', 'Salesforce Attribution', 'Monthly'],
        ['Event Pipeline Generated', '$5.8M', 'Salesforce Events Module', 'Post-event'],
    ]
    table4 = doc.add_table(rows=len(kpi_data), cols=4)
    table4.style = 'Table Grid'
    for i, row_data in enumerate(kpi_data):
        for j, val in enumerate(row_data):
            cell = table4.cell(i, j)
            run = cell.paragraphs[0].add_run(val)
            if i == 0:
                run.bold = True

    doc.add_heading('Approval Sign-Offs', level=2)
    doc.add_paragraph(
        'The following stakeholders must review and approve this document before campaign launch:'
    )
    approvals = [
        'CMO: Rachel Horne — Due by March 10, 2025',
        'CFO: Jonathan Blackwell — Due by March 14, 2025',
        'VP Engineering: Samantha Wu — Due by March 12, 2025',
        'Legal Counsel: David Fong — Due by March 13, 2025',
        'CEO: Marcus Delacroix — Final approval by March 17, 2025',
    ]
    for a in approvals:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_heading('Document Revision History', level=2)
    rev_data = [
        ['Version', 'Date', 'Author', 'Change Summary'],
        ['1.0', 'Feb 3, 2025', 'Rachel Horne', 'Initial draft'],
        ['1.5', 'Feb 14, 2025', 'Claire Okafor', 'Added event strategy'],
        ['2.0', 'Feb 24, 2025', 'Marcus team', 'Budget updated after CFO review'],
        ['2.4', 'Mar 1, 2025', 'Rachel Horne', 'Final pre-approval draft'],
    ]
    table5 = doc.add_table(rows=len(rev_data), cols=4)
    table5.style = 'Table Grid'
    for i, row_data in enumerate(rev_data):
        for j, val in enumerate(row_data):
            cell = table5.cell(i, j)
            run = cell.paragraphs[0].add_run(val)
            if i == 0:
                run.bold = True

    # ---- Footer: page number only (NO disclaimer) ----
    add_page_number_footer(section)

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
