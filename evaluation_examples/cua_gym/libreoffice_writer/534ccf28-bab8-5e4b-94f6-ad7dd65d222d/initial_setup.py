"""
Initial Setup: Talent Acquisition Strategy Document (narrative only)
Task ID: writer_hr_091
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_091'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('Talent Acquisition Strategy 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Global Human Resources Division')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    run.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Prepared by: Director of Talent Acquisition, Samantha Rivera')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_para2 = doc.add_paragraph()
    date_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para2.add_run('Date: January 15, 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacing

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines the comprehensive talent acquisition strategy for fiscal year 2026. '
        'As our organization continues to scale across North America, Europe, and Asia-Pacific markets, '
        'the demand for top-tier talent has intensified. Our hiring volume is projected to increase by 35% '
        'compared to 2025, requiring a fundamental shift in how we attract, evaluate, and onboard candidates.'
    )
    doc.add_paragraph(
        'Key strategic priorities include reducing our average time-to-fill from 42 days to 30 days, '
        'improving offer acceptance rates from 78% to 88%, and decreasing cost-per-hire by 15% through '
        'technology-enabled sourcing and employer branding initiatives. This strategy is built on data-driven '
        'insights from our 2025 recruitment analytics and industry benchmarking studies conducted by Mercer '
        'and LinkedIn Talent Solutions.'
    )
    doc.add_paragraph(
        'The following sections present our current hiring metrics, source effectiveness analysis, '
        'time-to-fill benchmarks by department, candidate pipeline optimization framework, employer branding '
        'action plan, technology stack recommendations, and a phased 12-month implementation roadmap.'
    )

    # --- Section 2: Current Hiring Metrics ---
    doc.add_heading('2. Current Hiring Metrics Dashboard', level=1)
    doc.add_paragraph(
        'The metrics below represent our organizational hiring performance as of Q4 2025. These figures '
        'serve as the baseline against which all 2026 targets will be measured. Data was collected from '
        'our existing ATS (Greenhouse), HRIS (Workday), and finance systems across all business units.'
    )
    doc.add_paragraph(
        'Total positions filled in 2025: 847. Open requisitions as of January 2026: 132. '
        'Average applications per role: 156. Internal mobility rate: 18%. Diversity hiring rate: 42%. '
        'Employee referral conversion rate: 34%. New hire 90-day retention: 91%.'
    )

    # --- Section 3: Source Effectiveness Analysis ---
    doc.add_heading('3. Source Effectiveness Analysis', level=1)
    doc.add_paragraph(
        'Understanding the effectiveness of each recruiting channel is critical for optimizing our '
        'talent acquisition budget of $4.2M. The analysis below evaluates 10 sourcing channels across '
        'four dimensions: total applications generated, interviews conducted, successful hires, and '
        'cost-per-hire. Data spans the full 2025 fiscal year.'
    )
    doc.add_paragraph(
        'Our top-performing channels by quality-of-hire metric continue to be Employee Referrals and '
        'LinkedIn Recruiter, while Indeed and company career page drive the highest application volumes. '
        'University partnerships and industry events show the highest cost-per-hire but contribute '
        'significantly to our early-career pipeline and employer brand visibility.'
    )

    # --- Section 4: Time-to-Fill Analysis ---
    doc.add_heading('4. Time-to-Fill Analysis by Department', level=1)
    doc.add_paragraph(
        'Time-to-fill varies significantly across departments due to differences in role complexity, '
        'market supply-demand dynamics, and hiring manager engagement. Engineering and Data Science '
        'roles consistently take the longest to fill due to fierce competition for specialized talent. '
        'Administrative and Customer Service roles benefit from larger candidate pools and faster '
        'screening processes.'
    )
    doc.add_paragraph(
        'Our 2026 targets aim to reduce time-to-fill across all departments by an average of 20%, '
        'primarily through pre-built talent pipelines, structured interviewing frameworks, and '
        'streamlined approval workflows. Departments exceeding 45-day averages will receive dedicated '
        'sourcing support from the talent acquisition team.'
    )

    # --- Section 5: Candidate Pipeline ---
    doc.add_heading('5. Candidate Pipeline Funnel Analysis', level=1)
    doc.add_paragraph(
        'The candidate pipeline funnel provides visibility into conversion rates at each stage of '
        'our recruitment process. Identifying bottlenecks and drop-off points allows us to implement '
        'targeted interventions. Our 2025 data reveals that the largest drop-off occurs between the '
        'phone screen and on-site interview stages, suggesting misalignment in initial screening criteria.'
    )
    doc.add_paragraph(
        'To address this, we are implementing structured phone screen rubrics aligned with hiring '
        'manager expectations and investing in recruiter training programs focused on competency-based '
        'assessment techniques.'
    )

    # --- Section 6: Employer Branding ---
    doc.add_heading('6. Employer Branding Action Plan', level=1)
    doc.add_paragraph(
        'Employer branding is a critical lever for attracting passive candidates and improving '
        'application quality. According to LinkedIn research, companies with strong employer brands '
        'see 50% more qualified applicants and reduce cost-per-hire by 43%. Our 2026 employer branding '
        'initiatives focus on content marketing, employee advocacy, social media presence, and '
        'candidate experience optimization.'
    )
    doc.add_paragraph(
        'We will partner with the Marketing Communications team to develop authentic employee stories, '
        'behind-the-scenes content, and thought leadership pieces. Additionally, we will launch a '
        'Glassdoor response program and redesign our careers page to highlight our culture, benefits, '
        'and growth opportunities.'
    )

    # --- Section 7: Technology Stack ---
    doc.add_heading('7. Technology Stack Comparison', level=1)
    doc.add_paragraph(
        'As our current ATS contract expires in Q2 2026, we have an opportunity to evaluate '
        'next-generation talent acquisition platforms. Three leading solutions have been shortlisted '
        'based on functionality, integration capabilities, scalability, and total cost of ownership: '
        'Greenhouse Advanced, Lever Enterprise, and SmartRecruiters CRM Pro.'
    )
    doc.add_paragraph(
        'Each platform has been assessed across eight feature categories by a cross-functional '
        'evaluation committee comprising Talent Acquisition, HR Technology, IT Security, and Finance '
        'representatives. The comparison below summarizes our findings.'
    )

    # --- Section 8: Implementation Roadmap ---
    doc.add_heading('8. 12-Month Implementation Roadmap', level=1)
    doc.add_paragraph(
        'The following roadmap outlines the phased rollout of all strategic initiatives described in '
        'this document. Each phase has been designed to build upon the previous one, ensuring sustainable '
        'change management and stakeholder alignment. Key milestones, owners, and success metrics are '
        'defined for each quarter.'
    )
    doc.add_paragraph(
        'Phase 1 (Q1) focuses on foundation-building: technology vendor selection, employer branding '
        'asset development, and recruiter training. Phase 2 (Q2) emphasizes process optimization and '
        'ATS migration. Phase 3 (Q3) targets scaling of new sourcing channels and analytics dashboards. '
        'Phase 4 (Q4) concentrates on performance evaluation and 2027 strategy formulation.'
    )

    # --- Conclusion ---
    doc.add_heading('9. Conclusion and Next Steps', level=1)
    doc.add_paragraph(
        'This strategy represents a data-driven, technology-enabled approach to talent acquisition '
        'that aligns with our organizational growth objectives. Successful execution depends on '
        'cross-functional collaboration between HR, Finance, IT, and business unit leaders. '
        'Monthly progress reviews will be conducted by the Talent Acquisition Steering Committee, '
        'with quarterly board updates on key performance indicators.'
    )
    doc.add_paragraph(
        'For questions or feedback regarding this strategy, please contact Samantha Rivera, '
        'Director of Talent Acquisition, at s.rivera@company.com or extension 4521.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
