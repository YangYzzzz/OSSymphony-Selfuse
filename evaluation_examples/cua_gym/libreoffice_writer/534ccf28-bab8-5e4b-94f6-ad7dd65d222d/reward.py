"""
Reward Script: Talent Acquisition Strategy Document with Data Visualizations
Task ID: writer_hr_091
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Document has >= 7 tables (task requires 7+ analytical tables)
  Component 2 (0.20): Source effectiveness table with 10 sourcing channels (11 rows, 6 cols)
  Component 3 (0.20): ATS comparison table with 3 systems across 8 features (9 rows, 4 cols)
  Component 4 (0.15): 12-month implementation roadmap table (>= 12 data rows)
  Component 5 (0.20): Colored section divider bars (>= 5 paragraph borders with non-default colors)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_091'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_tables = len(doc.tables)
    print(f"INFO: Document has {len(doc.paragraphs)} paragraphs, {num_tables} tables")

    # =========================================================================
    # Component 1: Document has >= 7 tables (0.25 points)
    # Task requires: hiring metrics dashboard tables, source effectiveness,
    # time-to-fill, candidate pipeline, employer branding, ATS comparison,
    # and implementation roadmap = at least 7 tables.
    # Initial has 0 tables, so this only passes on golden.
    # =========================================================================
    try:
        if num_tables >= 7:
            print(f"PASS: Component 1 — Document has {num_tables} tables (>= 7 required) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Document has {num_tables} tables, need >= 7")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Source effectiveness table with 10 sourcing channels (0.20 pts)
    # Must have: header row + 10 data rows = 11 rows, with columns for
    # applications, interviews, hires, cost-per-hire.
    # We look for a table with a header cell containing "Sourcing Channel"
    # or similar, with >= 11 rows and >= 5 columns.
    # =========================================================================
    try:
        source_table_found = False
        for t in doc.tables:
            # Check header row for sourcing-related keywords
            if len(t.rows) < 2:
                continue
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            header_text = ' '.join(header_cells)
            if ('sourcing' in header_text or 'source' in header_text or 'channel' in header_text) \
                    and len(t.columns) >= 5 and len(t.rows) >= 11:
                # Verify it has data about applications/interviews/hires
                has_apps = any('application' in h for h in header_cells)
                has_hires = any('hire' in h for h in header_cells)
                if has_apps or has_hires:
                    # Count non-empty data rows
                    data_rows = sum(1 for r in t.rows[1:] if r.cells[0].text.strip())
                    if data_rows >= 10:
                        source_table_found = True
                        print(f"PASS: Component 2 — Source effectiveness table found: {len(t.rows)} rows, {len(t.columns)} cols, {data_rows} sourcing channels (0.20 pts)")
                        total_score += 0.20
                        break
        if not source_table_found:
            print(f"FAIL: Component 2 — No source effectiveness table with 10+ channels found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: ATS comparison table — 3 systems across 8 features (0.20 pts)
    # Must have: header row with 3 ATS system names + feature column = 4 cols,
    # and 8 feature rows + header = 9 rows.
    # =========================================================================
    try:
        ats_table_found = False
        for t in doc.tables:
            if len(t.rows) < 2:
                continue
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            header_text = ' '.join(header_cells)
            # Look for "feature" in header and 4 columns (feature + 3 systems)
            if ('feature' in header_text) and len(t.columns) >= 4 and len(t.rows) >= 9:
                # Verify header has ATS-like system names (not just generic)
                non_feature_headers = [c.text.strip() for c in t.rows[0].cells[1:]]
                systems_count = sum(1 for h in non_feature_headers if len(h) > 3)
                if systems_count >= 3:
                    # Count feature rows (non-empty first column)
                    feature_rows = sum(1 for r in t.rows[1:] if r.cells[0].text.strip())
                    if feature_rows >= 8:
                        ats_table_found = True
                        print(f"PASS: Component 3 — ATS comparison table found: {len(t.rows)} rows, {len(t.columns)} cols, {feature_rows} features, systems: {non_feature_headers[:3]} (0.20 pts)")
                        total_score += 0.20
                        break
        if not ats_table_found:
            print(f"FAIL: Component 3 — No ATS comparison table (3 systems, 8+ features) found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================================
    # Component 4: 12-month implementation roadmap table (0.15 pts)
    # Must have >= 12 data rows (one per month) with columns for month,
    # initiative, milestone, owner, etc.
    # =========================================================================
    try:
        roadmap_found = False
        for t in doc.tables:
            if len(t.rows) < 2:
                continue
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            header_text = ' '.join(header_cells)
            # Look for month/timeline and milestone/initiative columns
            has_month = any(kw in header_text for kw in ['month', 'timeline', 'phase'])
            has_initiative = any(kw in header_text for kw in ['initiative', 'milestone', 'action', 'activity'])
            if has_month and has_initiative and len(t.rows) >= 13:
                data_rows = sum(1 for r in t.rows[1:] if r.cells[0].text.strip())
                if data_rows >= 12:
                    roadmap_found = True
                    print(f"PASS: Component 4 — Implementation roadmap table found: {data_rows} months (0.15 pts)")
                    total_score += 0.15
                    break
        if not roadmap_found:
            print(f"FAIL: Component 4 — No 12-month implementation roadmap table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =========================================================================
    # Component 5: Colored section divider bars (0.20 pts)
    # Task requires "section breaks with colored divider bars".
    # We check for paragraphs with bottom border elements that have
    # non-default colors. Need at least 5 such dividers.
    # Initial document has 0 dividers.
    # =========================================================================
    try:
        from docx.oxml.ns import qn
        divider_count = 0
        for p in doc.paragraphs:
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    for child in pBdr:
                        color = child.get(qn('w:color'))
                        if color and color.lower() not in ('auto', '000000', 'ffffff'):
                            divider_count += 1
                            break  # count one divider per paragraph

        if divider_count >= 5:
            print(f"PASS: Component 5 — Found {divider_count} colored section dividers (>= 5 required) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 5 — Found {divider_count} colored dividers, need >= 5")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
