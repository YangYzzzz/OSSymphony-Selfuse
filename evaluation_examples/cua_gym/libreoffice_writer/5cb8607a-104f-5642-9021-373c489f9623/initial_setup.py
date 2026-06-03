"""
Initial Setup: Bilingual reading exercise with two blocks of paragraphs
Task ID: wrpara_045
Domain: libreoffice_writer

Creates a document with 8 paragraphs:
  Paragraphs 1-4: English text (E1-E4)
  Paragraphs 5-8: Spanish translations (S1-S4)
Arranged as two sequential blocks (all English, then all Spanish).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# English paragraphs (E1-E4)
english_paragraphs = [
    "The Amazon rainforest, often referred to as the lungs of the Earth, spans across nine countries in South America. It produces approximately twenty percent of the world's oxygen and houses an estimated ten percent of all species on the planet. Deforestation and climate change pose significant threats to this vital ecosystem.",
    "Renewable energy sources such as solar, wind, and hydroelectric power have seen remarkable growth over the past decade. Governments worldwide are investing heavily in green infrastructure to reduce carbon emissions. By 2030, experts predict that renewables could account for nearly half of global electricity generation.",
    "The human brain contains roughly eighty-six billion neurons, each forming thousands of connections with neighboring cells. Modern neuroscience uses advanced imaging techniques like functional MRI to map these intricate networks. Understanding how the brain processes language and memory remains one of science's greatest challenges.",
    "Urban agriculture is transforming city landscapes around the world, from rooftop gardens in Tokyo to vertical farms in Singapore. These initiatives not only provide fresh produce to local communities but also reduce transportation costs and carbon footprints. Community-supported agriculture programs have doubled in participation since 2020.",
]

# Spanish paragraphs (S1-S4) - translations matching the English
spanish_paragraphs = [
    "La selva amazonica, a menudo conocida como los pulmones de la Tierra, se extiende por nueve paises de America del Sur. Produce aproximadamente el veinte por ciento del oxigeno mundial y alberga un estimado del diez por ciento de todas las especies del planeta. La deforestacion y el cambio climatico representan amenazas significativas para este ecosistema vital.",
    "Las fuentes de energia renovable como la solar, la eolica y la hidroelectrica han experimentado un crecimiento notable en la ultima decada. Los gobiernos de todo el mundo estan invirtiendo fuertemente en infraestructura verde para reducir las emisiones de carbono. Para 2030, los expertos predicen que las renovables podrian representar casi la mitad de la generacion electrica mundial.",
    "El cerebro humano contiene aproximadamente ochenta y seis mil millones de neuronas, cada una formando miles de conexiones con las celulas vecinas. La neurociencia moderna utiliza tecnicas avanzadas de imagen como la resonancia magnetica funcional para mapear estas intrincadas redes. Comprender como el cerebro procesa el lenguaje y la memoria sigue siendo uno de los mayores desafios de la ciencia.",
    "La agricultura urbana esta transformando los paisajes de las ciudades en todo el mundo, desde jardines en azoteas en Tokio hasta granjas verticales en Singapur. Estas iniciativas no solo proporcionan productos frescos a las comunidades locales, sino que tambien reducen los costos de transporte y la huella de carbono. Los programas de agricultura apoyada por la comunidad han duplicado su participacion desde 2020.",
]


def create_initial():
    doc = Document()

    # Add a title
    title = doc.add_heading("Bilingual Reading Exercise", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add English block (paragraphs 1-4)
    for i, text in enumerate(english_paragraphs):
        para = doc.add_paragraph(text)
        # Default spacing, no special formatting

    # Add Spanish block (paragraphs 5-8)
    for i, text in enumerate(spanish_paragraphs):
        para = doc.add_paragraph(text)
        # Default spacing, no special formatting

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
