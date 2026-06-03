"""
Reward Script: Biology glossary color-coding by word length
Task ID: osworld_writer_colorword_005
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4 pts): All 4-or-fewer-char words are colored teal (#008080)
  Component 2 (0.4 pts): All 5-or-more-char words are colored brown (#8B4513)
  Component 3 (0.2 pts): All 45 cells have been colored (none left in default black)

Ground truth (from context):
  - 9 rows x 5 cols table = 45 cells total
  - 4 or fewer chars → teal (RGB 0x00, 0x80, 0x80 = #008080)
  - 5 or more chars → brown (RGB 0x8B, 0x45, 0x13 = #8B4513)
  - All 45 cells must be colored
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_005'

# Target colors
TEAL = (0x00, 0x80, 0x80)    # #008080 — for words with 4 or fewer characters
BROWN = (0x8B, 0x45, 0x13)   # #8B4513 — for words with 5 or more characters
BLACK = (0x00, 0x00, 0x00)   # #000000 — initial state (should score 0)

# Tolerance for color matching (Euclidean distance in RGB space)
COLOR_TOLERANCE = 40


def color_distance(rgb_obj, target_tuple):
    """Compute Euclidean distance between an RGBColor and an (R, G, B) tuple."""
    r, g, b = rgb_obj[0], rgb_obj[1], rgb_obj[2]
    return sqrt((r - target_tuple[0])**2 + (g - target_tuple[1])**2 + (b - target_tuple[2])**2)


def get_cell_color(cell):
    """
    Extract the text color of the first run in a table cell.
    Returns an RGBColor object or None if no color is set (default/inherited).
    """
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.color and run.font.color.rgb is not None:
                return run.font.color.rgb
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    # Load document — if this fails, score is 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Verify table exists and has expected structure
    if not doc.tables:
        print("CRITICAL: No tables found in document.")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    if num_rows != 9 or num_cols != 5:
        print(f"CRITICAL: Expected 9x5 table, found {num_rows}x{num_cols}.")
        print("REWARD: 0.0")
        return 0.0

    # Collect all cell data
    cells = []
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text:
                color = get_cell_color(cell)
                cells.append((text, color))

    total_cells = len(cells)
    if total_cells == 0:
        print("CRITICAL: No text cells found in table.")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: All 4-or-fewer-char words are colored teal (0.4 points)
    # These FAIL on initial_env (all black) and PASS on golden_env (teal applied)
    # -------------------------------------------------------------------------
    try:
        short_words = [(text, color) for (text, color) in cells if len(text) <= 4]
        short_teal = 0
        short_wrong = []

        for text, color in short_words:
            if color is not None and color_distance(color, BLACK) > COLOR_TOLERANCE:
                # Some color has been applied — check if it's teal
                if color_distance(color, TEAL) < COLOR_TOLERANCE:
                    short_teal += 1
                else:
                    short_wrong.append((text, str(color)))
            else:
                # Still black/default — not colored
                short_wrong.append((text, "black/default"))

        expected_short = len(short_words)
        if short_teal == expected_short and expected_short > 0:
            print(f"PASS: Component 1 — all {expected_short} short words (≤4 chars) are teal (#008080) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — {short_teal}/{expected_short} short words are teal; incorrect: {short_wrong[:5]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: All 5-or-more-char words are colored brown (0.4 points)
    # These FAIL on initial_env (all black) and PASS on golden_env (brown applied)
    # -------------------------------------------------------------------------
    try:
        long_words = [(text, color) for (text, color) in cells if len(text) >= 5]
        long_brown = 0
        long_wrong = []

        for text, color in long_words:
            if color is not None and color_distance(color, BLACK) > COLOR_TOLERANCE:
                # Some color has been applied — check if it's brown
                if color_distance(color, BROWN) < COLOR_TOLERANCE:
                    long_brown += 1
                else:
                    long_wrong.append((text, str(color)))
            else:
                long_wrong.append((text, "black/default"))

        expected_long = len(long_words)
        if long_brown == expected_long and expected_long > 0:
            print(f"PASS: Component 2 — all {expected_long} long words (≥5 chars) are brown (#8B4513) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — {long_brown}/{expected_long} long words are brown; incorrect: {long_wrong[:5]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: All 45 cells have been colored (none in default black) (0.2 pts)
    # This FAILS on initial_env (all black) and PASSES on golden_env (all colored)
    # -------------------------------------------------------------------------
    try:
        uncolored = []
        for text, color in cells:
            if color is None or color_distance(color, BLACK) <= COLOR_TOLERANCE:
                uncolored.append(text)

        if len(uncolored) == 0:
            print(f"PASS: Component 3 — all {total_cells} cells have been colored (none in default black) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — {len(uncolored)} cells still in default black: {uncolored[:5]}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against the canonical file path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
