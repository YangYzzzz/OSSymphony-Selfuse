"""
Initial Setup: Biology glossary table with 45 terms in 9x5 table, all black text
Task ID: osworld_writer_colorword_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_colorword_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title paragraph
    title = doc.add_heading("Biology Glossary", level=1)

    # Subtitle
    subtitle = doc.add_paragraph(
        "Key biological terms and concepts for introductory biology students."
    )

    # 45 biology terms arranged in a 9x5 table
    # Mix of short (<=4 chars) and long (>=5 chars) terms
    terms = [
        # Row 1
        'cell', 'nucleus', 'DNA', 'mitosis', 'gene',
        # Row 2
        'RNA', 'meiosis', 'ATP', 'cytoplasm', 'chloroplast',
        # Row 3
        'acid', 'membrane', 'base', 'ribosome', 'vacuole',
        # Row 4
        'vein', 'enzyme', 'ions', 'protein', 'bacteria',
        # Row 5
        'stem', 'flagella', 'root', 'osmosis', 'pore',
        # Row 6
        'leaf', 'glucose', 'bone', 'hormone', 'nerve',
        # Row 7
        'skin', 'plasma', 'lung', 'insulin', 'blood',
        # Row 8
        'gill', 'neuron', 'bile', 'synapse', 'fern',
        # Row 9
        'moss', 'embryo', 'seed', 'photosynthesis', 'xylem',
    ]

    # Verify we have exactly 45 terms
    assert len(terms) == 45, f"Expected 45 terms, got {len(terms)}"

    # Create a 9-row x 5-column table
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'

    # Populate table cells with terms, all in default black text
    term_idx = 0
    for row_idx in range(9):
        for col_idx in range(5):
            cell = table.cell(row_idx, col_idx)
            # Clear existing paragraph and add new run with black text
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(terms[term_idx])
            run.font.size = Pt(11)
            # Explicitly set black color (default, no color-coding)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            term_idx += 1

    # Add a footer note
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "Note: Color-coding is an accessibility feature to help students distinguish term length."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
