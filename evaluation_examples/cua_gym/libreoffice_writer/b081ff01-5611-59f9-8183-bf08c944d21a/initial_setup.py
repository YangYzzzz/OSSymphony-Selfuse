"""
Initial Setup: Create a report document with TOC that has two erroneously styled paragraphs
Task ID: writer_mt_094
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_094'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Filler paragraphs to bulk up page count - realistic report body text
FILLER_PARAGRAPHS = [
    "The fiscal year showed significant improvements across multiple performance indicators. Revenue growth exceeded projections by approximately 12%, driven primarily by expansion into new market segments and improved customer retention rates.",
    "Our analysis of quarterly trends reveals a steady upward trajectory in both gross margins and operating efficiency. Cost optimization initiatives launched in Q1 yielded measurable results by mid-year, with procurement savings alone contributing to a 3.2% reduction in overall operating expenses.",
    "Customer satisfaction metrics, as measured by our annual Net Promoter Score survey, improved from 47 to 62 points year-over-year. This improvement correlates strongly with investments made in customer support infrastructure and the deployment of our new self-service portal.",
    "Employee engagement scores remained stable at 78%, with notable improvements in the areas of professional development opportunities and workplace flexibility. The hybrid work model, now in its second full year, continues to receive positive feedback from both management and staff.",
    "Technology investments during the period totaled $4.2 million, primarily allocated toward cloud infrastructure migration, cybersecurity enhancements, and the development of proprietary analytics tools. These investments are expected to generate returns within a 24-month horizon.",
    "Supply chain resilience was tested during Q3 when regional disruptions affected three of our primary suppliers. Contingency protocols activated successfully, limiting production delays to fewer than five business days across all product lines.",
    "The marketing team executed 14 major campaigns across digital and traditional channels. Digital advertising spend increased by 28% while maintaining a customer acquisition cost below our target threshold of $45 per qualified lead.",
    "Research and development efforts concentrated on three strategic priorities: product miniaturization, energy efficiency improvements, and integration of machine learning capabilities into existing product offerings.",
    "International operations expanded to two additional markets, bringing the total to 23 countries. Currency fluctuations had a modest negative impact on reported revenues, accounting for approximately $1.8 million in translation losses.",
    "The board of directors approved a revised capital allocation framework that increases the dividend payout ratio from 30% to 35% while maintaining sufficient reserves for strategic acquisitions and organic growth initiatives.",
    "Regulatory compliance remained a key focus area, with the legal team successfully navigating new data privacy regulations in the European Union and updated environmental standards in North America.",
    "Partnership and alliance activities accelerated during the second half of the year. Three new strategic partnerships were formalized, providing access to complementary technologies and distribution channels.",
    "Quality assurance metrics improved across all manufacturing facilities. Defect rates decreased from 2.1% to 1.4%, and on-time delivery performance reached a record high of 97.3%.",
    "The sustainability program achieved several milestones, including a 15% reduction in carbon emissions per unit produced and the attainment of ISO 14001 certification at two additional facilities.",
    "Looking ahead, management has identified five strategic priorities for the coming fiscal year: digital transformation acceleration, talent acquisition in emerging technology domains, geographic expansion in Southeast Asia, product portfolio rationalization, and deepening customer engagement through personalized experiences.",
    "Financial projections for the next fiscal year anticipate revenue growth in the range of 8-12%, supported by the full-year contribution of recently launched products and continued momentum in subscription-based service offerings.",
    "Risk management frameworks were updated to reflect evolving geopolitical considerations and emerging cybersecurity threats. Scenario planning exercises conducted in Q4 informed the development of enhanced contingency protocols.",
    "The annual audit, completed without material findings, confirmed the integrity of financial reporting processes and the effectiveness of internal controls over financial reporting.",
]


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc(doc):
    """Add a Table of Contents field to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr_text = run2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr_text)

    run3 = paragraph.add_run()
    fld_char_separate = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run3._element.append(fld_char_separate)

    # Placeholder TOC entries (will be updated by LibreOffice)
    run4 = paragraph.add_run("[Table of Contents - Right-click and Update to refresh]")

    run5 = paragraph.add_run()
    fld_char_end = run5._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run5._element.append(fld_char_end)


def add_filler(doc, count=3, start_idx=0):
    """Add filler paragraphs to pad page content."""
    for i in range(count):
        idx = (start_idx + i) % len(FILLER_PARAGRAPHS)
        doc.add_paragraph(FILLER_PARAGRAPHS[idx])


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('Annual Performance Report 2024', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by the Strategic Planning Division')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('March 2025')
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_page_break()

    # Table of Contents page
    toc_heading = doc.add_heading('Table of Contents', level=0)
    add_toc(doc)
    doc.add_page_break()

    # --- Section 1: Executive Summary (Heading 1) ---
    doc.add_heading('Executive Summary', level=1)
    add_filler(doc, count=4, start_idx=0)
    doc.add_page_break()

    # --- Section 2: Financial Overview (Heading 1) ---
    doc.add_heading('Financial Overview', level=1)
    add_filler(doc, count=3, start_idx=4)

    # Sub-section: Revenue Analysis (Heading 2)
    doc.add_heading('Revenue Analysis', level=2)
    add_filler(doc, count=3, start_idx=7)
    doc.add_page_break()

    # Sub-section: Cost Structure (Heading 2)
    doc.add_heading('Cost Structure', level=2)
    add_filler(doc, count=3, start_idx=10)

    # --- ERRONEOUS PARAGRAPH 1 (page ~8): Styled as Heading 2 but should be Body Text ---
    doc.add_heading('Note: This data is preliminary', level=2)
    add_filler(doc, count=3, start_idx=13)
    doc.add_page_break()

    # Sub-section: Profit Margins (Heading 2)
    doc.add_heading('Profit Margins', level=2)
    add_filler(doc, count=4, start_idx=16)
    doc.add_page_break()

    # --- Section 3: Operational Performance (Heading 1) ---
    doc.add_heading('Operational Performance', level=1)
    add_filler(doc, count=3, start_idx=1)

    # Sub-section: Manufacturing Efficiency (Heading 2)
    doc.add_heading('Manufacturing Efficiency', level=2)
    add_filler(doc, count=4, start_idx=5)
    doc.add_page_break()

    # Sub-section: Supply Chain Management (Heading 2)
    doc.add_heading('Supply Chain Management', level=2)
    add_filler(doc, count=4, start_idx=9)
    doc.add_page_break()

    # --- Section 4: Market Analysis (Heading 1) ---
    doc.add_heading('Market Analysis', level=1)
    add_filler(doc, count=3, start_idx=2)

    # Sub-section: Competitive Landscape (Heading 2)
    doc.add_heading('Competitive Landscape', level=2)
    add_filler(doc, count=4, start_idx=6)
    doc.add_page_break()

    # --- Section 5: Human Resources (Heading 1) ---
    doc.add_heading('Human Resources', level=1)
    add_filler(doc, count=3, start_idx=3)

    # Sub-section: Talent Acquisition (Heading 2)
    doc.add_heading('Talent Acquisition', level=2)
    add_filler(doc, count=4, start_idx=8)

    # --- ERRONEOUS PARAGRAPH 2 (page ~15): Styled as Heading 2 but should be Body Text ---
    doc.add_heading('Source: Annual Survey 2024', level=2)
    add_filler(doc, count=3, start_idx=12)
    doc.add_page_break()

    # --- Section 6: Strategic Outlook (Heading 1) ---
    doc.add_heading('Strategic Outlook', level=1)
    add_filler(doc, count=4, start_idx=14)

    # Sub-section: Risk Assessment (Heading 2)
    doc.add_heading('Risk Assessment', level=2)
    add_filler(doc, count=3, start_idx=0)

    # Conclusion
    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)
    add_filler(doc, count=3, start_idx=15)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count headings for verification
    heading_count = 0
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and para.style.name != 'Heading':
            heading_count += 1
    print(f'Total headings in document (including erroneous): {heading_count}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
