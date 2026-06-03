"""
Initial Setup: Research essay with placeholder for citation insertion
Task ID: osworld_writer_biblio_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_biblio_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Advances in Natural Language Processing: A Survey', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author line
    author_para = doc.add_paragraph('Zhang Wei, Department of Computer Science, University of Technology')
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # blank line

    # --- Abstract Heading ---
    doc.add_heading('Abstract', level=2)

    # --- Paragraph 1: Introduction ---
    para1 = doc.add_paragraph(
        'Natural language processing (NLP) has undergone remarkable transformation over the past decade, '
        'driven largely by advances in machine learning and the availability of large-scale training datasets. '
        'Early rule-based systems have been supplanted by statistical models, which in turn have given way to '
        'neural network architectures that achieve state-of-the-art performance across a wide range of tasks. '
        'This survey examines the evolution of NLP techniques and highlights the most significant breakthroughs '
        'that have shaped the field.'
    )

    # --- Paragraph 2: Contains [REF] placeholder ---
    para2 = doc.add_paragraph(
        'The introduction of transformer-based models has been particularly transformative, enabling unprecedented '
        'performance gains in machine translation, sentiment analysis, and question answering. Recent studies have '
        'demonstrated that pre-trained language models can be fine-tuned for domain-specific applications with '
        'relatively modest amounts of labeled data [REF]. These findings suggest that transfer learning will '
        'continue to play a central role in the development of practical NLP systems for both commercial and '
        'research applications in the coming years.'
    )

    # --- Paragraph 3: Methodology ---
    doc.add_heading('Methodology', level=2)
    para3 = doc.add_paragraph(
        'This review synthesizes findings from over 120 peer-reviewed publications from 2018 to 2024. '
        'Papers were selected based on their citation count, novelty of approach, and relevance to core NLP '
        'tasks including named entity recognition, coreference resolution, and semantic parsing. '
        'We conducted systematic database searches across ACL Anthology, IEEE Xplore, and Google Scholar, '
        'applying inclusion criteria focused on empirical results and reproducibility.'
    )

    # --- Paragraph 4: Results ---
    doc.add_heading('Results and Discussion', level=2)
    para4 = doc.add_paragraph(
        'Our analysis reveals three dominant paradigm shifts in the NLP literature. First, the migration '
        'from recurrent neural networks to self-attention mechanisms substantially improved both training '
        'efficiency and downstream task performance. Second, the emergence of instruction-tuned language models '
        'has blurred the boundary between supervised and zero-shot learning. Third, multilingual models '
        'trained on diverse corpora have closed the performance gap between high-resource and low-resource '
        'languages, with notable gains in cross-lingual transfer benchmarks such as XTREME and XGLUE.'
    )

    # --- Paragraph 5: Conclusion ---
    doc.add_heading('Conclusion', level=2)
    para5 = doc.add_paragraph(
        'The trajectory of NLP research points toward increasingly generalized models capable of handling '
        'diverse linguistic tasks without task-specific engineering. Future work should address outstanding '
        'challenges in reasoning under uncertainty, grounded language understanding, and efficient inference '
        'for resource-constrained environments. The interdisciplinary nature of these problems will require '
        'collaboration across linguistics, cognitive science, and systems engineering.'
    )

    # --- References Section ---
    doc.add_heading('References', level=2)

    ref1 = doc.add_paragraph(
        '1. Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., & '
        'Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, '
        '30, 5998-6008.'
    )

    ref2 = doc.add_paragraph(
        '2. Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep '
        'bidirectional transformers for language understanding. Proceedings of NAACL-HLT 2019, 4171-4186. '
        'https://doi.org/10.18653/v1/N19-1423'
    )

    ref3 = doc.add_paragraph(
        '3. Brown, T. B., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., Neelakantan, A., '
        'Shyam, P., Sastry, G., Askell, A., & Amodei, D. (2020). Language models are few-shot learners. '
        'Advances in Neural Information Processing Systems, 33, 1877-1901.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
