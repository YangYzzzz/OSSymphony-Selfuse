"""
Initial Setup: Configure page break before Chapter 2 heading
Task ID: writer_fs_006
Domain: libreoffice_writer

Creates a multi-chapter research document where Chapter 2: Methods
follows directly after Chapter 1 text with no page break.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Research Report: Sustainable Urban Agriculture', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by Dr. Elena Vasquez and Dr. Tomasz Kowalski')
    run.font.size = Pt(12)
    run.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('March 2025')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Chapter 1: Introduction ---
    ch1 = doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_paragraph(
        'Urban agriculture has emerged as a critical component of sustainable city planning '
        'in the 21st century. As global populations increasingly concentrate in metropolitan '
        'areas, the need for localized food production systems has become more pressing than '
        'ever before. This report examines the current state of urban farming initiatives '
        'across twelve major cities in North America and Europe.'
    )

    doc.add_paragraph(
        'The concept of growing food within city boundaries is not new. Historical records '
        'from World War II show that victory gardens in the United States produced up to 40% '
        'of the nation\'s vegetables. However, modern urban agriculture differs significantly '
        'in both scale and technological sophistication. Vertical farming, hydroponic systems, '
        'and rooftop greenhouses now supplement traditional community garden plots.'
    )

    doc.add_heading('1.1 Background', level=2)

    doc.add_paragraph(
        'Between 2018 and 2024, investment in urban agriculture technology grew from $1.2 billion '
        'to $8.7 billion globally. Major players including AeroFarms, Gotham Greens, and Plenty '
        'have demonstrated that indoor vertical farms can achieve yields 350 times greater per '
        'square foot than conventional agriculture. Meanwhile, cities like Detroit, Singapore, '
        'and Copenhagen have integrated food production into their municipal planning frameworks.'
    )

    doc.add_paragraph(
        'Despite these advances, significant challenges remain. High energy costs, limited crop '
        'diversity in controlled environments, and the socioeconomic barriers to community '
        'participation continue to hinder widespread adoption. This study aims to quantify these '
        'barriers and propose evidence-based solutions drawn from successful case studies.'
    )

    doc.add_heading('1.2 Research Objectives', level=2)

    doc.add_paragraph(
        'The primary objectives of this research are threefold: (1) to evaluate the economic '
        'viability of urban farming models across different climate zones, (2) to assess the '
        'nutritional impact on food-insecure neighborhoods, and (3) to develop a scalable '
        'framework for municipal integration of urban agriculture programs.'
    )

    # --- Chapter 2: Methods (NO page break - this is the task target) ---
    ch2 = doc.add_heading('Chapter 2: Methods', level=1)

    doc.add_paragraph(
        'This study employed a mixed-methods research design combining quantitative yield data '
        'analysis with qualitative interviews of urban farmers and municipal planners. Data '
        'collection spanned 18 months from January 2023 through June 2024.'
    )

    doc.add_heading('2.1 Study Sites', level=2)

    doc.add_paragraph(
        'Twelve cities were selected based on population density, existing urban agriculture '
        'infrastructure, and climate diversity: New York, Chicago, Detroit, San Francisco, '
        'Toronto, London, Amsterdam, Copenhagen, Berlin, Singapore, Tokyo, and Melbourne. '
        'Within each city, between three and seven farming sites were identified for detailed '
        'analysis, yielding a total of 58 study locations.'
    )

    doc.add_heading('2.2 Data Collection', level=2)

    doc.add_paragraph(
        'Quantitative data included monthly crop yields measured in kilograms per square meter, '
        'energy consumption in kilowatt-hours, water usage in liters, and operational costs in '
        'local currency normalized to USD. Qualitative data were gathered through semi-structured '
        'interviews with 142 stakeholders including farm managers, city officials, nutritionists, '
        'and community organizers.'
    )

    # --- Chapter 3: Results ---
    ch3 = doc.add_heading('Chapter 3: Results', level=1)

    doc.add_paragraph(
        'Analysis of the collected data revealed significant variation in urban farming '
        'productivity and economic sustainability across the twelve study cities. Indoor vertical '
        'farms consistently outperformed rooftop and community garden models in terms of yield '
        'per unit area, though at substantially higher capital and operating costs.'
    )

    doc.add_heading('3.1 Yield Comparisons', level=2)

    doc.add_paragraph(
        'Average leafy green yields in vertical farms reached 42.3 kg/m\u00b2 annually, compared '
        'to 8.7 kg/m\u00b2 for rooftop gardens and 5.2 kg/m\u00b2 for community plots. However, '
        'when adjusted for energy input, community gardens demonstrated the highest efficiency '
        'ratio at 3.8 kg per kWh equivalent, versus 0.9 kg/kWh for vertical farms.'
    )

    doc.add_paragraph(
        'The Detroit case study proved particularly instructive. The city\'s extensive vacant '
        'lot inventory, combined with municipal tax incentives introduced in 2021, enabled a '
        'network of 23 community farms to supply fresh produce to over 15,000 residents in '
        'previously designated food deserts. The program achieved a cost-benefit ratio of 2.4:1 '
        'within its first three years of operation.'
    )

    # --- Chapter 4: Discussion ---
    ch4 = doc.add_heading('Chapter 4: Discussion', level=1)

    doc.add_paragraph(
        'The findings of this study underscore the importance of context-specific approaches '
        'to urban agriculture implementation. No single model emerged as universally superior; '
        'rather, the most successful programs were those that aligned farming methodology with '
        'local climate conditions, available real estate, community needs, and municipal policy '
        'support structures.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
