"""
Reward Script: Configure 'Chapter 2: Methods' heading with page break before
Task ID: writer_fs_006
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Chapter 2 paragraph has page_break_before enabled
  Component 2 (0.3): page_break_before is set AND no manual page breaks used (correct method)
  Component 3 (0.2): page_break_before is set AND only on Chapter 2 (targeted application)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_006'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the 'Chapter 2: Methods' paragraph
    chapter2_para = None
    chapter2_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('Chapter 2'):
            chapter2_para = para
            chapter2_idx = i
            break

    if chapter2_para is None:
        print("CRITICAL: Could not find 'Chapter 2: Methods' paragraph")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found 'Chapter 2: Methods' at paragraph index {chapter2_idx}")

    # Check the core requirement once — used by all components
    pbb_is_true = False
    try:
        pbb = chapter2_para.paragraph_format.page_break_before
        pbb_is_true = (pbb is True)
        print(f"INFO: page_break_before = {pbb}")
    except Exception as e:
        print(f"ERROR: Could not read page_break_before: {e}")

    # Component 1: Chapter 2 paragraph has page_break_before enabled (0.5 points)
    # This is the core task requirement.
    try:
        if pbb_is_true:
            print(f"PASS: Component 1 — page_break_before is True on Chapter 2 (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — page_break_before is not True")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: page_break_before is set AND no manual page breaks near Chapter 2 (0.3 points)
    # The task specifies using paragraph style page break, NOT a manual break.
    # This component is anchored to the task change: it only awards points if
    # page_break_before is True AND no manual breaks were used instead.
    try:
        if not pbb_is_true:
            print(f"FAIL: Component 2 — page_break_before not set (prerequisite)")
        else:
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            manual_break_found = False

            # Check the Chapter 2 paragraph itself for manual breaks
            for run in chapter2_para.runs:
                for br in run.element.findall('.//w:br', ns):
                    btype = br.attrib.get(qn('w:type'), 'line')
                    if btype == 'page':
                        manual_break_found = True

            # Also check the paragraph immediately before Chapter 2
            if chapter2_idx > 0:
                prev_para = doc.paragraphs[chapter2_idx - 1]
                for run in prev_para.runs:
                    for br in run.element.findall('.//w:br', ns):
                        btype = br.attrib.get(qn('w:type'), 'line')
                        if btype == 'page':
                            manual_break_found = True

            if not manual_break_found:
                print(f"PASS: Component 2 — page_break_before set with no manual breaks (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Manual page break detected alongside paragraph break")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: page_break_before set AND only on Chapter 2 among Heading 1 paras (0.2 points)
    # Confirms the change was targeted to Chapter 2 only, not a blanket style modification.
    # Anchored to the task change: only awards points if page_break_before is True.
    try:
        if not pbb_is_true:
            print(f"FAIL: Component 3 — page_break_before not set (prerequisite)")
        else:
            other_heading1_with_pbb = []
            for i, para in enumerate(doc.paragraphs):
                if i == chapter2_idx:
                    continue
                if para.style and para.style.name == 'Heading 1':
                    if para.paragraph_format.page_break_before is True:
                        other_heading1_with_pbb.append(para.text[:50])

            if len(other_heading1_with_pbb) == 0:
                print(f"PASS: Component 3 — page_break_before only on Chapter 2, not other headings (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Other headings also have page_break_before: {other_heading1_with_pbb}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
