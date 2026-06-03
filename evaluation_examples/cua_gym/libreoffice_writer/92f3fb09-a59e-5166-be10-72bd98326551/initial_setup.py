"""
Initial Setup: Create a Writer document with editorial article containing five underlined key phrases.
Task ID: writer_rd_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_053'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

KEY_PHRASES = [
    'critical issue',
    'unprecedented growth',
    'immediate action required',
    'fundamental change',
    'long-term strategy',
]

# Article paragraphs with placeholders for key phrases marked by <<phrase>>
ARTICLE_PARAGRAPHS = [
    (
        "The Future of Sustainable Urban Development",
        "heading",
    ),
    (
        "In recent years, city planners across the globe have faced a <<critical issue>> "
        "that demands attention from policymakers, architects, and community leaders alike. "
        "The rapid pace of urbanization has strained infrastructure, housing markets, and "
        "public services in ways that few anticipated just a decade ago. Metropolitan areas "
        "that once thrived on manufacturing are now transitioning to knowledge-based economies, "
        "bringing both opportunities and significant challenges.",
        "body",
    ),
    (
        "According to the latest report from the International Urban Planning Consortium, "
        "cities in Southeast Asia and sub-Saharan Africa have experienced <<unprecedented growth>> "
        "over the past five years. Population densities in cities such as Jakarta, Lagos, and "
        "Dhaka have increased by nearly thirty percent, placing enormous pressure on water supply "
        "networks, transportation corridors, and waste management systems. The report highlights "
        "that without coordinated intervention, these trends will accelerate beyond the capacity "
        "of local governments to respond effectively.",
        "body",
    ),
    (
        "Municipal leaders in several pilot cities have already declared that <<immediate action required>> "
        "is the guiding principle behind their new sustainability frameworks. In Barcelona, for example, "
        "the Superblock initiative has reclaimed entire city blocks from vehicular traffic, transforming "
        "them into pedestrian plazas with greenery, outdoor seating, and community gathering spaces. "
        "Early results show a fourteen percent reduction in nitrogen dioxide levels and a measurable "
        "improvement in residents' reported quality of life.",
        "body",
    ),
    (
        "Experts argue that what we are witnessing is nothing short of a <<fundamental change>> "
        "in how cities conceptualize their relationship with the natural environment. Traditional "
        "planning models that prioritized automobile access and single-use zoning are giving way to "
        "mixed-use developments, transit-oriented design, and green infrastructure investments. "
        "Professor Elena Vasquez of the Urban Resilience Institute notes that this shift reflects "
        "a broader societal recognition that environmental health and economic prosperity are not "
        "competing goals but deeply intertwined outcomes.",
        "body",
    ),
    (
        "Looking ahead, the consortium recommends that every metropolitan region develop a "
        "<<long-term strategy>> spanning at least twenty-five years. Such a strategy should "
        "incorporate climate adaptation measures, affordable housing targets, renewable energy "
        "integration, and digital infrastructure expansion. The report emphasizes that short-term "
        "fixes, while sometimes necessary, cannot substitute for comprehensive planning that "
        "anticipates demographic shifts, technological disruption, and evolving community needs.",
        "body",
    ),
    (
        "As cities around the world grapple with these complex realities, the lessons emerging "
        "from early adopters offer both caution and hope. The path toward sustainable urban "
        "development is neither simple nor predetermined, but the growing consensus among "
        "researchers, practitioners, and civic leaders suggests that meaningful progress is "
        "within reach for communities willing to invest in bold, evidence-based approaches.",
        "body",
    ),
]


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_font(run, font_name='Liberation Serif', font_size=11, bold=False, underline=False):
    """Apply standard body formatting to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.underline = underline


def add_paragraph_with_phrases(doc, text, key_phrases):
    """Add a paragraph, splitting on key phrases to underline them."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.paragraph_format.space_after = Pt(8)

    # Split text around <<phrase>> markers
    remaining = text
    while '<<' in remaining:
        before, rest = remaining.split('<<', 1)
        phrase, remaining = rest.split('>>', 1)

        # Add text before the phrase
        if before:
            run = para.add_run(before)
            set_run_font(run)

        # Add the key phrase with underline
        run = para.add_run(phrase)
        set_run_font(run, underline=True)

    # Add any remaining text
    if remaining:
        run = para.add_run(remaining)
        set_run_font(run)

    return para


def create_initial():
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style.font.size = Pt(11)

    for text, para_type in ARTICLE_PARAGRAPHS:
        if para_type == 'heading':
            heading = doc.add_heading(text, level=1)
            for run in heading.runs:
                run.font.name = 'Liberation Serif'
        elif '<<' in text:
            add_paragraph_with_phrases(doc, text, KEY_PHRASES)
        else:
            para = doc.add_paragraph(text)
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.paragraph_format.space_after = Pt(8)
            for run in para.runs:
                set_run_font(run)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
