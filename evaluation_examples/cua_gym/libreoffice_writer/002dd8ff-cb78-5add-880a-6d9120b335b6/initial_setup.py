"""
Initial Setup: Certificate template document with no page borders
Task ID: writer_fs_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard letter size with generous margins for certificate look
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # NO page borders - this is what the task requires the agent to add

    # --- Certificate Title ---
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("CERTIFICATE OF ACHIEVEMENT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)  # Navy blue
    run.font.name = "Times New Roman"

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("This certificate is proudly presented to")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Recipient Name ---
    name_para = doc.add_paragraph()
    name_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_para.paragraph_format.space_before = Pt(12)
    name_para.paragraph_format.space_after = Pt(24)
    run = name_para.add_run("Dr. Elena Vasquez-Morales")
    run.bold = True
    run.font.size = Pt(36)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)

    # --- Achievement Description ---
    desc = doc.add_paragraph()
    desc.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    desc.paragraph_format.space_after = Pt(6)
    run = desc.add_run("in recognition of outstanding contributions to the")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    field = doc.add_paragraph()
    field.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    field.paragraph_format.space_after = Pt(6)
    run = field.add_run("International Conference on Computational Linguistics")
    run.bold = True
    run.italic = True
    run.font.size = Pt(15)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)

    detail = doc.add_paragraph()
    detail.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    detail.paragraph_format.space_after = Pt(36)
    run = detail.add_run(
        "for the paper entitled \"Bridging Semantic Gaps in Multilingual "
        "Transformer Architectures\" presented on March 15, 2025 in Geneva, Switzerland"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_after = Pt(36)
    run = date_para.add_run("Awarded on March 17, 2025")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.italic = True

    # --- Signature Block ---
    sig1 = doc.add_paragraph()
    sig1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig1.paragraph_format.space_before = Pt(24)
    sig1.paragraph_format.space_after = Pt(0)
    run = sig1.add_run("_________________________________")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    sig_name1 = doc.add_paragraph()
    sig_name1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig_name1.paragraph_format.space_after = Pt(0)
    run = sig_name1.add_run("Prof. Richard Thornton")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    sig_title1 = doc.add_paragraph()
    sig_title1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig_title1.paragraph_format.space_after = Pt(18)
    run = sig_title1.add_run("Conference Chair, ICCL 2025")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.italic = True

    sig2 = doc.add_paragraph()
    sig2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig2.paragraph_format.space_after = Pt(0)
    run = sig2.add_run("_________________________________")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    sig_name2 = doc.add_paragraph()
    sig_name2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sig_name2.paragraph_format.space_after = Pt(0)
    run = sig_name2.add_run("Dr. Aisha Patel")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    sig_title2 = doc.add_paragraph()
    sig_title2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sig_title2.add_run("Secretary General, Association for Computational Linguistics")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
