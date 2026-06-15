"""
Initial Setup: Create a Writer document with technical content but no custom macros.
Task ID: writer_tech_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # -- Title --
    title = doc.add_heading("API Integration Guide", level=0)

    # -- Subtitle / intro --
    intro = doc.add_paragraph()
    run = intro.add_run("Version 2.4 — Last updated March 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    intro.paragraph_format.space_after = Pt(12)

    # -- Section 1 --
    doc.add_heading("1. Authentication", level=1)

    p1 = doc.add_paragraph()
    p1.add_run("All API requests require a valid bearer token. Use the ")
    r1 = p1.add_run("generateToken()")
    r1.font.name = "Liberation Mono"
    r1.font.size = Pt(10)
    p1.add_run(" method from the authentication module to obtain a session token. "
               "Tokens expire after 3600 seconds by default.")

    p2 = doc.add_paragraph()
    p2.add_run("The token should be included in the ")
    r2 = p2.add_run("Authorization")
    r2.font.name = "Liberation Mono"
    r2.font.size = Pt(10)
    p2.add_run(" header of every HTTP request. For example:")

    # Code-like block paragraph
    code_block = doc.add_paragraph()
    code_block.paragraph_format.left_indent = Inches(0.5)
    cb_run = code_block.add_run('curl -H "Authorization: Bearer $TOKEN" https://api.example.com/v2/users')
    cb_run.font.name = "Liberation Mono"
    cb_run.font.size = Pt(9)

    # -- Section 2 --
    doc.add_heading("2. Rate Limiting", level=1)

    p3 = doc.add_paragraph(
        "The API enforces rate limits of 100 requests per minute per API key. "
        "When the limit is exceeded, the server responds with HTTP status 429 "
        "and includes a Retry-After header indicating the number of seconds to wait."
    )

    p4 = doc.add_paragraph()
    p4.add_run("Implement exponential backoff in your client using ")
    r3 = p4.add_run("retryWithBackoff(maxRetries=5)")
    r3.font.name = "Liberation Mono"
    r3.font.size = Pt(10)
    p4.add_run(". This function handles transient failures and respects the server's "
               "rate limit headers automatically.")

    # -- Section 3 --
    doc.add_heading("3. Endpoint Reference", level=1)

    # Table of endpoints
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    headers = ["Method", "Endpoint", "Description", "Auth Required"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    endpoints = [
        ["GET", "/v2/users", "List all users with pagination", "Yes"],
        ["POST", "/v2/users", "Create a new user account", "Yes"],
        ["GET", "/v2/users/{id}", "Retrieve user by ID", "Yes"],
        ["PUT", "/v2/users/{id}/profile", "Update user profile fields", "Yes"],
        ["DELETE", "/v2/users/{id}", "Deactivate user account", "Admin"],
    ]
    for r, row_data in enumerate(endpoints, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if c == 1:
                run.font.name = "Liberation Mono"

    doc.add_paragraph()  # spacer

    # -- Section 4 --
    doc.add_heading("4. Error Handling", level=1)

    p5 = doc.add_paragraph(
        "All error responses follow a consistent JSON structure with three fields: "
        "error_code (integer), message (string), and details (object, optional). "
        "Your client should inspect the error_code field to determine the appropriate "
        "recovery action."
    )

    p6 = doc.add_paragraph()
    p6.add_run("Common error codes include ")
    codes = [("AUTH_EXPIRED", ", "), ("RATE_LIMITED", ", "),
             ("RESOURCE_NOT_FOUND", ", "), ("VALIDATION_ERROR", ".")]
    for code_text, suffix in codes:
        r = p6.add_run(code_text)
        r.font.name = "Liberation Mono"
        r.font.size = Pt(10)
        p6.add_run(suffix)

    p7 = doc.add_paragraph()
    p7.add_run("Wrap all API calls in a try-catch block using the ")
    r4 = p7.add_run("ApiErrorHandler")
    r4.font.name = "Liberation Mono"
    r4.font.size = Pt(10)
    p7.add_run(" class, which provides centralized logging and metrics collection "
               "for failed requests.")

    # -- Section 5 --
    doc.add_heading("5. Webhook Configuration", level=1)

    p8 = doc.add_paragraph(
        "Webhooks allow your application to receive real-time notifications when "
        "specific events occur. Configure webhook endpoints through the dashboard "
        "or programmatically via the API."
    )

    doc.add_paragraph("Supported webhook events:", style="List Bullet")
    events = [
        "user.created — Triggered when a new user registers",
        "user.updated — Triggered when profile data changes",
        "payment.completed — Triggered after successful payment processing",
        "subscription.cancelled — Triggered when a subscription ends",
        "export.ready — Triggered when a data export file is available",
    ]
    for event in events:
        doc.add_paragraph(event, style="List Bullet")

    p9 = doc.add_paragraph()
    p9.add_run("Each webhook payload includes a ")
    r5 = p9.add_run("X-Webhook-Signature")
    r5.font.name = "Liberation Mono"
    r5.font.size = Pt(10)
    p9.add_run(" header for verification. Validate this signature using HMAC-SHA256 "
               "with your webhook secret before processing any payload data.")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
