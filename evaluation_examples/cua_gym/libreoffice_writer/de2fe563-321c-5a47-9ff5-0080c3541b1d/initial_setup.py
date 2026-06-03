"""
Initial Setup: Mail Merge Wizard — Purchase Order Template with Suppliers data source
Task ID: writer_mt_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import csv
import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DB_DIR = f'{WORKDIR}/suppliers_db'
CSV_FILE = f'{DB_DIR}/suppliers.csv'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_suppliers_csv():
    """Create the Suppliers CSV data source."""
    os.makedirs(DB_DIR, exist_ok=True)
    rows = [
        ['SupplierName', 'ContactPerson', 'Address', 'City', 'State', 'Zip', 'PaymentTerms'],
        ['Apex Industrial Supply', 'Maria Gonzalez', '4500 Commerce Blvd', 'Dallas', 'TX', '75201', 'Net 30'],
        ['Northern Components Ltd', 'James Chen', '891 Industrial Parkway', 'Toronto', 'ON', 'M5V 2T6', 'Net 45'],
        ['Pacific Rim Materials', 'Yuki Tanaka', '2200 Harbor Drive', 'Long Beach', 'CA', '90802', 'Net 60'],
        ['Summit Hardware Co', 'Robert Andersen', '156 Mountain View Rd', 'Denver', 'CO', '80202', 'Net 30'],
        ['EuroTech Fasteners GmbH', 'Hans Mueller', '78 Industriestrasse', 'Stuttgart', 'BW', '70173', '2/10 Net 30'],
        ['Coastal Packaging Inc', 'Sarah Williams', '3400 Bayshore Blvd', 'Tampa', 'FL', '33629', 'Net 45'],
        ['Midwest Steel Solutions', 'David Kowalski', '900 Foundry Lane', 'Pittsburgh', 'PA', '15201', 'Net 30'],
    ]
    with open(CSV_FILE, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f'Suppliers CSV created: {CSV_FILE}')


def register_datasource():
    """Register the 'Suppliers' data source in LibreOffice."""
    # LibreOffice stores registered data sources in registereddb/datasources.xml
    lo_profile = os.path.expanduser('~/.config/libreoffice/4/user')
    reg_dir = os.path.join(lo_profile, 'registereddb')
    os.makedirs(reg_dir, exist_ok=True)

    reg_file = os.path.join(reg_dir, 'datasources.xml')

    # Create the registration XML
    xml_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<office:data-sources xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0">
</office:data-sources>
'''
    # We also need to create the .odb file for the data source
    odb_path = f'{DB_DIR}/Suppliers.odb'
    create_odb_for_csv(odb_path)

    # Register via LibreOffice macro using command line
    # Actually, the simplest way is to use the datasources.xcu configuration
    xcu_dir = os.path.join(lo_profile, 'registrymodifications.xcu')

    # Better approach: create an .odb file and register it via the xcu config
    # LibreOffice reads data source registrations from:
    # ~/.config/libreoffice/4/user/registrymodifications.xcu

    # Read existing xcu if present
    if os.path.exists(xcu_dir):
        with open(xcu_dir, 'r') as f:
            xcu_content = f.read()
    else:
        xcu_content = '''<?xml version="1.0" encoding="UTF-8"?>
<oor:items xmlns:oor="http://openoffice.org/2001/registry"
           xmlns:xs="http://www.w3.org/2001/XMLSchema"
           xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
</oor:items>
'''

    # Add data source registration entry if not already there
    if 'Suppliers' not in xcu_content:
        registration_entry = f'''<item oor:path="/org.openoffice.Office.DataAccess/RegisteredNames"><node oor:name="Suppliers" oor:op="replace"><prop oor:name="Location" oor:op="fuse"><value>file://{odb_path}</value></prop><prop oor:name="Name" oor:op="fuse"><value>Suppliers</value></prop></node></item>'''

        # Insert before closing tag
        xcu_content = xcu_content.replace('</oor:items>', registration_entry + '\n</oor:items>')

        with open(xcu_dir, 'w') as f:
            f.write(xcu_content)
        print(f'Registered Suppliers data source in {xcu_dir}')


def create_odb_for_csv(odb_path):
    """Create a minimal .odb file that references the CSV directory."""
    import zipfile
    import io

    # An .odb is a ZIP with specific XML files inside
    content_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
    xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    xmlns:db="urn:oasis:names:tc:opendocument:xmlns:database:1.0"
    office:version="1.2">
  <office:body>
    <office:database>
      <db:data-source>
        <db:connection-data>
          <db:connection-resource db:href="sdbc:flat:{DB_DIR}"/>
        </db:connection-data>
        <db:driver-settings db:system-driver-settings="" db:base-dn="" db:parameter-name-substitution="false">
          <db:auto-increment/>
          <db:delimiter db:field="," db:string="&quot;" db:decimal="." db:thousand=""/>
          <db:character-set db:encoding="UTF-8"/>
        </db:driver-settings>
        <db:application-connection-settings db:is-table-name-length-limited="false"
            db:append-table-alias-name="false" db:max-row-count="100">
          <db:table-filter>
            <db:table-include-filter>
              <db:table-filter-pattern>%</db:table-filter-pattern>
            </db:table-include-filter>
          </db:table-filter>
        </db:application-connection-settings>
      </db:data-source>
    </office:database>
  </office:body>
</office:document-content>'''

    meta_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta
    xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
    office:version="1.2">
  <office:meta/>
</office:document-meta>'''

    manifest_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">
  <manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.database"/>
  <manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>
  <manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/>
</manifest:manifest>'''

    mimetype = 'application/vnd.oasis.opendocument.database'

    with zipfile.ZipFile(odb_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        # mimetype must be first and uncompressed
        zf.writestr('mimetype', mimetype, compress_type=zipfile.ZIP_STORED)
        zf.writestr('content.xml', content_xml)
        zf.writestr('meta.xml', meta_xml)
        zf.writestr('META-INF/manifest.xml', manifest_xml)

    print(f'ODB file created: {odb_path}')


def create_po_template():
    """Create the Purchase Order template document."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Company header
    header_para = doc.add_paragraph()
    header_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.paragraph_format.space_after = Pt(2)
    run = header_para.add_run('PINNACLE MANUFACTURING INC.')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Company address
    addr_para = doc.add_paragraph()
    addr_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr_para.paragraph_format.space_after = Pt(0)
    run = addr_para.add_run('8200 Innovation Drive, Suite 300')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    addr2_para = doc.add_paragraph()
    addr2_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr2_para.paragraph_format.space_after = Pt(0)
    run = addr2_para.add_run('Austin, TX 78701')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_para.paragraph_format.space_after = Pt(6)
    run = contact_para.add_run('Phone: (512) 555-0147  |  Fax: (512) 555-0148  |  procurement@pinnaclemfg.com')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Horizontal line
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(4)
    line_para.paragraph_format.space_after = Pt(12)
    run = line_para.add_run('_' * 72)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    run.font.size = Pt(10)

    # PO Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(16)
    run = title_para.add_run('PURCHASE ORDER')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # PO number and date
    info_para = doc.add_paragraph()
    info_para.paragraph_format.space_after = Pt(2)
    run = info_para.add_run('PO Number: ')
    run.bold = True
    run.font.size = Pt(11)
    run2 = info_para.add_run('PO-2026-00487')
    run2.font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(2)
    run = date_para.add_run('Date: ')
    run.bold = True
    run.font.size = Pt(11)
    run2 = date_para.add_run('April 2, 2026')
    run2.font.size = Pt(11)

    terms_para = doc.add_paragraph()
    terms_para.paragraph_format.space_after = Pt(16)
    run = terms_para.add_run('Payment Terms: ')
    run.bold = True
    run.font.size = Pt(11)
    run2 = terms_para.add_run('Per supplier agreement')
    run2.font.size = Pt(11)

    # Supplier section (placeholder - this is what mail merge will fill)
    supplier_heading = doc.add_paragraph()
    supplier_heading.paragraph_format.space_after = Pt(4)
    run = supplier_heading.add_run('SHIP TO / SUPPLIER:')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Placeholder lines for supplier info (to be replaced by merge fields)
    placeholder_lines = [
        '[Supplier Name]',
        '[Contact Person]',
        '[Street Address]',
        '[City, State ZIP]',
    ]
    for line in placeholder_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Blank line
    doc.add_paragraph()

    # Salutation placeholder
    sal_para = doc.add_paragraph()
    sal_para.paragraph_format.space_after = Pt(8)
    run = sal_para.add_run('Dear [Contact Person],')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Body text
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(8)
    run = body1.add_run(
        'Please find below the details of our purchase order. We kindly request '
        'that you confirm receipt of this order and provide an estimated delivery '
        'date within 5 business days.'
    )
    run.font.size = Pt(11)

    # Order table
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'

    headers = ['Item #', 'Description', 'Qty', 'Unit Price', 'Total']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    items = [
        ['1', 'Stainless Steel Hex Bolts M10x50 (Box/100)', '25', '$42.50', '$1,062.50'],
        ['2', 'Industrial Grade Silicone Gaskets 6"', '150', '$8.75', '$1,312.50'],
        ['3', 'Precision Ball Bearings 6205-2RS', '80', '$15.20', '$1,216.00'],
        ['4', 'Heat-Resistant Ceramic Washers M12', '200', '$3.40', '$680.00'],
    ]
    for r, row_data in enumerate(items, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val
            for run in table.cell(r, c).paragraphs[0].runs:
                run.font.size = Pt(10)

    # Blank line after table
    doc.add_paragraph()

    # Total line
    total_para = doc.add_paragraph()
    total_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = total_para.add_run('Order Total: $4,271.00')
    run.bold = True
    run.font.size = Pt(12)

    # Closing
    doc.add_paragraph()
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.space_after = Pt(4)
    run = closing_para.add_run(
        'If you have any questions regarding this purchase order, please contact '
        'our Procurement Department at (512) 555-0147 or procurement@pinnaclemfg.com.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    thanks_para = doc.add_paragraph()
    thanks_para.paragraph_format.space_after = Pt(4)
    run = thanks_para.add_run('Thank you for your continued partnership.')
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_after = Pt(0)
    run = sig_para.add_run('Jennifer Blackwell')
    run.bold = True
    run.font.size = Pt(11)

    title_sig = doc.add_paragraph()
    title_sig.paragraph_format.space_after = Pt(0)
    run = title_sig.add_run('Director of Procurement')
    run.font.size = Pt(11)

    company_sig = doc.add_paragraph()
    run = company_sig.add_run('Pinnacle Manufacturing Inc.')
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'PO template created: {OUTPUT}')


def main():
    # Create the Suppliers CSV data source
    create_suppliers_csv()

    # Create the ODB and register the data source
    register_datasource()

    # Create the PO template document
    create_po_template()

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
