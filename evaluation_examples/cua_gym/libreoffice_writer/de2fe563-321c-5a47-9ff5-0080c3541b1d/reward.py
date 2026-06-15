"""
Reward Script: Mail Merge Wizard - Form Letter with Suppliers Data Source
Task ID: writer_mt_041
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Address block has merge fields replacing placeholders
  Component 2 (0.25): Salutation line uses ContactPerson merge field
  Component 3 (0.25): All expected merge field names present (from Suppliers data source)
  Component 4 (0.20): Document structure integrity preserved
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_041'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the Mail Merge Wizard was completed:
    - Placeholder text replaced with MERGEFIELD codes
    - Address block fields present
    - Salutation field present
    - Correct field names from Suppliers data source
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get the full body XML for merge field detection
    try:
        body_xml = doc.element.xml
    except Exception as e:
        print(f"CRITICAL: Cannot read document XML: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract all MERGEFIELD names from instrText elements
    merge_field_names = re.findall(r'MERGEFIELD\s+(\w+)', body_xml)
    merge_field_set = set(merge_field_names)

    # Component 1: Address block has merge fields replacing placeholders (0.30 points)
    # In the initial file, paragraphs 10-13 have placeholder text like [Supplier Name].
    # In the golden file, these are replaced with MERGEFIELD codes.
    # We check that the old placeholders are gone AND merge fields exist in address area.
    try:
        # Check that placeholder text is no longer present
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        has_placeholder_supplier = '[Supplier Name]' in all_text
        has_placeholder_address = '[Street Address]' in all_text
        has_placeholder_city = '[City, State ZIP]' in all_text

        # Check that address-related merge fields exist
        address_fields = {'SupplierName', 'Address', 'City', 'State', 'Zip'}
        found_address_fields = address_fields.intersection(merge_field_set)

        if not has_placeholder_supplier and not has_placeholder_address and not has_placeholder_city and len(found_address_fields) >= 3:
            print(f"PASS: Component 1 — Address block has merge fields ({found_address_fields}), placeholders removed (0.30 pts)")
            total_score += 0.30
        else:
            missing = address_fields - found_address_fields
            print(f"FAIL: Component 1 — Placeholders still present: supplier={has_placeholder_supplier}, address={has_placeholder_address}, city={has_placeholder_city}. Missing fields: {missing}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Salutation line uses ContactPerson merge field (0.25 points)
    # Initial has "Dear [Contact Person]," — golden should have "Dear <<ContactPerson>>,"
    # The merge field renders as chevron characters in para.text
    try:
        salutation_merge_count = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            # Check for "Dear" followed by merge field indicator (chevrons or MERGEFIELD)
            if text.lower().startswith('dear'):
                # Check if this paragraph contains a ContactPerson merge field
                # The rendered text shows guillemets: Dear <<ContactPerson>>,
                if '\u00ab' in text and 'ContactPerson' in text:
                    salutation_merge_count += 1
                    break
                # Also check XML of this paragraph for MERGEFIELD ContactPerson
                para_xml = p._element.xml
                if 'MERGEFIELD' in para_xml and 'ContactPerson' in para_xml:
                    salutation_merge_count += 1
                    break

        # Also verify the old placeholder is gone
        has_placeholder_contact_salutation = 'Dear [Contact Person]' in all_text

        if salutation_merge_count > 0 and not has_placeholder_contact_salutation:
            print(f"PASS: Component 2 — Salutation uses ContactPerson merge field (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Salutation merge field count: {salutation_merge_count}, old placeholder present: {has_placeholder_contact_salutation}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All expected merge field names present from Suppliers data source (0.25 points)
    # The Suppliers data source has: SupplierName, ContactPerson, Address, City, State, Zip, PaymentTerms
    # The task requires address block and salutation, so we expect at minimum:
    # SupplierName, ContactPerson, Address, City, State, Zip
    try:
        expected_fields = {'SupplierName', 'ContactPerson', 'Address', 'City', 'State', 'Zip'}
        found_fields = expected_fields.intersection(merge_field_set)
        coverage = len(found_fields) / len(expected_fields)

        if coverage >= 1.0:
            print(f"PASS: Component 3 — All {len(expected_fields)} expected merge fields present: {found_fields} (0.25 pts)")
            total_score += 0.25
        elif coverage >= 0.5:
            partial = round(0.25 * coverage, 2)
            print(f"PARTIAL: Component 3 — {len(found_fields)}/{len(expected_fields)} fields present: {found_fields}, missing: {expected_fields - found_fields} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {len(found_fields)}/{len(expected_fields)} fields present: {found_fields}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Document structure integrity preserved (0.20 points)
    # The original template content should be preserved around the merge fields:
    # - Company header "PINNACLE MANUFACTURING INC."
    # - "PURCHASE ORDER" title
    # - PO Number, Date
    # - Order details table
    # - Closing signature
    try:
        integrity_checks = 0
        total_integrity = 4

        if 'PINNACLE MANUFACTURING INC.' in all_text:
            integrity_checks += 1
        if 'PURCHASE ORDER' in all_text:
            integrity_checks += 1
        if 'PO-2026-00487' in all_text:
            integrity_checks += 1
        # Check table still exists with order items
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            if len(table.rows) >= 2:
                integrity_checks += 1

        # Only award points if ALL integrity checks pass AND at least one merge field exists
        # (This ensures we don't score integrity on the initial file since it has no merge fields)
        if integrity_checks == total_integrity and len(merge_field_names) > 0:
            print(f"PASS: Component 4 — Document structure intact with merge fields present ({integrity_checks}/{total_integrity} checks) (0.20 pts)")
            total_score += 0.20
        elif integrity_checks == total_integrity and len(merge_field_names) == 0:
            print(f"FAIL: Component 4 — Document structure intact but no merge fields present (score conditional on merge fields)")
        else:
            print(f"FAIL: Component 4 — Document structure partially damaged ({integrity_checks}/{total_integrity} checks)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification
persist_app_state("libreoffice_writer")

# Run verification
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
