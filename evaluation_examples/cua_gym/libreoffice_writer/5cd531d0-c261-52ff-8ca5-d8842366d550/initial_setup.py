"""
Initial Setup: Company newsletter with single-column layout on all pages.
Task ID: writer_biz_029
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup for default section (single column, standard margins) --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # =====================================================
    # PAGE 1 - Newsletter title and introduction
    # =====================================================

    # Newsletter Title
    title = doc.add_heading('Meridian Monthly', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(36)

    # Subtitle / Edition line
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle.add_run('Volume 12, Issue 3  |  March 2026')
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub_run.italic = True

    # Horizontal rule (via bottom border on an empty paragraph)
    hr = doc.add_paragraph()
    hr_fmt = hr._element
    pBdr = hr_fmt.makeelement(qn('w:pBdr'), {})
    bottom = hr_fmt.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:space'): '1',
        qn('w:color'): '1F497D'
    })
    pBdr.append(bottom)
    pPr = hr_fmt.find(qn('w:pPr'))
    if pPr is None:
        pPr = hr_fmt.makeelement(qn('w:pPr'), {})
        hr_fmt.insert(0, pPr)
    pPr.append(pBdr)

    # Introduction paragraph
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(12)
    intro.paragraph_format.space_after = Pt(6)
    intro_run = intro.add_run(
        'Welcome to the March edition of Meridian Monthly, your trusted source for company news, '
        'industry insights, and team highlights. This month, we celebrate the successful launch of '
        'our new product line and recognize the outstanding contributions of our colleagues across '
        'all departments. As we enter the second quarter, we look forward to building on the '
        'momentum of an exceptional first quarter.'
    )
    intro_run.font.size = Pt(11)

    # A second intro paragraph to fill out page 1
    intro2 = doc.add_paragraph()
    intro2.paragraph_format.space_after = Pt(6)
    intro2_run = intro2.add_run(
        'In this issue, you will find updates on our regional expansion plans, a spotlight on the '
        'engineering team\'s recent hackathon, a recap of the annual charity gala, and practical tips '
        'from our wellness program. We also introduce a new column dedicated to professional development '
        'resources recommended by our leadership team. We hope you enjoy reading as much as we enjoyed '
        'putting it together. If you have story ideas or feedback, please reach out to the communications '
        'team at newsletter@meridian.com.'
    )
    intro2_run.font.size = Pt(11)

    # =====================================================
    # PAGE 2+ - Article content (still single column in initial)
    # =====================================================
    # Insert a section break (new page) so page 2 starts a new section
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.left_margin = Inches(1)
    new_section.right_margin = Inches(1)
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)

    # Article 1
    h1 = doc.add_heading('Regional Expansion: New Offices in Austin and Denver', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    p1_run = p1.add_run(
        'Meridian Corporation is thrilled to announce the opening of two new regional offices in '
        'Austin, Texas, and Denver, Colorado. The Austin office, located in the thriving East Side '
        'tech corridor, will serve as a hub for our growing software development team. Meanwhile, '
        'the Denver location will house our expanding sales and customer success operations for '
        'the Mountain West region.'
    )
    p1_run.font.size = Pt(10)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2_run = p2.add_run(
        '"These new offices represent a significant investment in our people and our future," said '
        'CEO Patricia Langford during the announcement. "We chose Austin and Denver because of their '
        'vibrant talent pools and the quality of life they offer our employees. We expect to hire '
        'approximately 120 new team members across both locations by the end of the year."'
    )
    p2_run.font.size = Pt(10)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    p3_run = p3.add_run(
        'The Austin office will open its doors on April 15, with a grand opening celebration planned '
        'for the following week. The Denver office is scheduled to begin operations on May 1. Both '
        'locations will feature modern open-plan workspaces, dedicated collaboration zones, a wellness '
        'room, and a fully equipped kitchen. Employees interested in transferring to either location '
        'should contact HR by March 31.'
    )
    p3_run.font.size = Pt(10)

    # Article 2
    h2 = doc.add_heading('Engineering Hackathon Yields Innovative Prototypes', level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    p4_run = p4.add_run(
        'Last month, the engineering department hosted its third annual hackathon, bringing together '
        '45 developers, designers, and product managers for 48 hours of intensive innovation. The '
        'event produced eight working prototypes, three of which have been greenlit for further '
        'development by the product steering committee.'
    )
    p4_run.font.size = Pt(10)

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(6)
    p5_run = p5.add_run(
        'The winning project, "SmartSync," developed by a cross-functional team led by senior engineer '
        'Raj Patel, demonstrated a real-time data synchronization framework that could reduce '
        'server load by up to 40%. "We had been thinking about this problem for months, but the '
        'hackathon gave us the focused time to actually build a proof of concept," Raj explained.'
    )
    p5_run.font.size = Pt(10)

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(6)
    p6_run = p6.add_run(
        'Runner-up projects included "GreenDash," an energy consumption monitoring dashboard for '
        'our data centers, and "ClientPulse," a sentiment analysis tool for customer support tickets. '
        'The hackathon committee has recommended expanding next year\'s event to include participants '
        'from all departments, recognizing that diverse perspectives lead to stronger solutions.'
    )
    p6_run.font.size = Pt(10)

    # Article 3
    h3 = doc.add_heading('Annual Charity Gala Raises Record Funds', level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(6)
    p7_run = p7.add_run(
        'The Meridian Foundation\'s annual charity gala, held at the Grand Pavilion on February 22, '
        'raised a record-breaking $285,000 for local education initiatives. The event, themed '
        '"Building Bridges to Tomorrow," attracted over 350 guests, including employees, community '
        'leaders, and partner organizations.'
    )
    p7_run.font.size = Pt(10)

    p8 = doc.add_paragraph()
    p8.paragraph_format.space_after = Pt(6)
    p8_run = p8.add_run(
        'Highlights of the evening included a keynote address by Dr. Amara Sullivan, a renowned '
        'education researcher, and a live auction featuring experiences donated by local businesses. '
        'The most popular auction item was a private cooking class with Chef Marco Benedetti, which '
        'sold for $12,500. All proceeds will fund scholarships, STEM program grants, and after-school '
        'tutoring services for underserved communities in the metropolitan area.'
    )
    p8_run.font.size = Pt(10)

    # Article 4
    h4 = doc.add_heading('Wellness Corner: Spring Into Healthy Habits', level=1)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p9 = doc.add_paragraph()
    p9.paragraph_format.space_after = Pt(6)
    p9_run = p9.add_run(
        'As the days grow longer and temperatures rise, our corporate wellness team encourages '
        'everyone to take advantage of the season by establishing healthier routines. Research shows '
        'that small, consistent changes are more effective than dramatic overhauls. Here are five '
        'simple habits to adopt this spring: take a 15-minute walk during lunch, swap one processed '
        'snack for a piece of fruit daily, drink an extra glass of water each afternoon, stretch for '
        'five minutes before starting work, and try a new physical activity once a week.'
    )
    p9_run.font.size = Pt(10)

    p10 = doc.add_paragraph()
    p10.paragraph_format.space_after = Pt(6)
    p10_run = p10.add_run(
        'The wellness team is also launching a company-wide step challenge starting April 1. Teams '
        'of four can register through the intranet portal. The winning team will receive wellness '
        'gift baskets and an extra personal day. Last year\'s challenge saw participation from over '
        '60% of employees and resulted in a measurable increase in overall energy and satisfaction '
        'scores in the quarterly engagement survey.'
    )
    p10_run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
