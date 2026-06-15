"""
Initial Setup: Create a Masters Thesis document with chapters, sub-sections, and figures
but NO Table of Contents or Table of Figures.
Task ID: writer_mt_092
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_092'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph, fmt='decimal'):
    """Add a PAGE field code to a paragraph. fmt: 'decimal' or 'lowerRoman'."""
    run1 = paragraph.add_run()
    fld_char_begin = run1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run1._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = run2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = ' PAGE '
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fld_char_end)


def set_page_number_format(section, fmt='lowerRoman', start=None):
    """Set page number format for a section via XML.
    fmt: 'lowerRoman', 'decimal', 'upperRoman', etc.
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = sectPr.makeelement(qn('w:pgNumType'), {})
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))


def create_initial():
    doc = Document()

    # -- Default style setup --
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # =============================================
    # SECTION 1: Front Matter (Roman numeral pages)
    # =============================================
    # The first section is the default section
    section1 = doc.sections[0]
    section1.page_width = Inches(8.5)
    section1.page_height = Inches(11)
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)
    section1.left_margin = Inches(1.25)
    section1.right_margin = Inches(1.25)

    # Set Roman numeral page numbering starting at i
    set_page_number_format(section1, fmt='lowerRoman', start=1)

    # Add footer with page number
    footer1 = section1.footer
    footer1.is_linked_to_previous = False
    fp = footer1.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp, fmt='lowerRoman')

    # -- Title Page (page i) --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    run = title_para.add_run('Machine Learning Approaches to\nNatural Language Understanding')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(48)
    run = subtitle.add_run('A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nMaster of Science in Computer Science')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(48)
    run = author.add_run('by\nElena Rodriguez')
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    university = doc.add_paragraph()
    university.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    university.paragraph_format.space_before = Pt(36)
    run = university.add_run('Department of Computer Science\nStanford University\nMarch 2025')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # -- Page break after title (page ii - reserved for TOC) --
    doc.add_page_break()

    # Page ii placeholder (where TOC should go - currently blank)
    placeholder_ii = doc.add_paragraph()
    placeholder_ii.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    placeholder_ii.paragraph_format.space_before = Pt(200)
    run = placeholder_ii.add_run('[This page reserved for Table of Contents]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # -- Page break (page iii - possible TOC continuation) --
    doc.add_page_break()

    placeholder_iii = doc.add_paragraph()
    placeholder_iii.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    placeholder_iii.paragraph_format.space_before = Pt(200)
    run = placeholder_iii.add_run('[This page intentionally left blank]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # -- Page break (page iv - reserved for Table of Figures) --
    doc.add_page_break()

    placeholder_iv = doc.add_paragraph()
    placeholder_iv.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    placeholder_iv.paragraph_format.space_before = Pt(200)
    run = placeholder_iv.add_run('[This page reserved for Table of Figures]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # =============================================
    # SECTION 2: Body (Arabic numeral pages)
    # =============================================
    # New section with page break
    new_section_para = doc.add_paragraph()
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_number_format(new_section, fmt='decimal', start=1)

    # Copy page dimensions
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)
    new_section.left_margin = Inches(1.25)
    new_section.right_margin = Inches(1.25)

    # Footer with page number
    footer2 = new_section.footer
    footer2.is_linked_to_previous = False
    fp2 = footer2.paragraphs[0]
    fp2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp2, fmt='decimal')

    # Chapter and section data
    chapters = [
        {
            'title': 'Introduction',
            'sections': [
                ('Background and Motivation', [
                    'The rapid advancement of natural language processing (NLP) has transformed how computers interact with human language. Over the past decade, deep learning architectures have achieved remarkable progress in tasks ranging from machine translation to sentiment analysis, fundamentally changing the landscape of computational linguistics.',
                    'This thesis investigates novel approaches to natural language understanding (NLU) that leverage transformer-based architectures and multi-task learning frameworks. Our work builds upon the foundational contributions of Vaswani et al. (2017) and extends them to domain-specific applications in biomedical text mining.',
                ]),
                ('Research Questions', [
                    'This research addresses three primary questions: (1) How can multi-task learning improve performance on low-resource NLU tasks? (2) What architectural modifications to transformer models best capture domain-specific linguistic patterns? (3) How does curriculum learning affect convergence in fine-tuning large language models?',
                ]),
                ('Thesis Organization', [
                    'The remainder of this thesis is organized as follows. Chapter 2 reviews related work in natural language understanding and transfer learning. Chapter 3 describes our proposed methodology. Chapters 4 and 5 present our experimental setup and results. Chapter 6 discusses implications, and Chapter 7 concludes with future directions.',
                ]),
            ],
            'figures': [
                ('Overview of the NLU pipeline architecture showing data flow from raw text through preprocessing, embedding, and classification stages', 1),
            ],
        },
        {
            'title': 'Literature Review',
            'sections': [
                ('Transformer Architectures', [
                    'The transformer architecture, introduced by Vaswani et al. (2017), revolutionized sequence modeling by replacing recurrent connections with self-attention mechanisms. The multi-head attention mechanism allows the model to jointly attend to information from different representation subspaces at different positions.',
                    'Subsequent work by Devlin et al. (2019) introduced BERT, which demonstrated that bidirectional pre-training of transformers on large unlabeled corpora could dramatically improve performance on a wide range of NLP benchmarks.',
                ]),
                ('Transfer Learning in NLP', [
                    'Transfer learning has become a cornerstone of modern NLP. Pre-training on large corpora followed by task-specific fine-tuning has shown consistent improvements across tasks. Howard and Ruder (2018) introduced ULMFiT, establishing key techniques such as discriminative fine-tuning and slanted triangular learning rates.',
                ]),
                ('Multi-task Learning', [
                    'Multi-task learning (MTL) aims to improve generalization by leveraging shared representations across related tasks. Caruana (1997) provided early theoretical foundations, while recent work by Liu et al. (2019) demonstrated effective multi-task approaches for NLU using shared transformer encoders.',
                ]),
                ('Domain Adaptation', [
                    'Domain adaptation techniques address the challenge of applying models trained on one domain to another. Gururangan et al. (2020) showed that continued pre-training on domain-specific text, termed domain-adaptive pre-training (DAPT), significantly improves downstream task performance in specialized domains such as biomedical and computer science text.',
                ]),
            ],
            'figures': [
                ('Comparison of transformer attention patterns across different layers for biomedical text versus general domain text', 2),
                ('Timeline of major NLP model developments from 2017 to 2024 showing parameter counts and benchmark scores', 3),
            ],
        },
        {
            'title': 'Methodology',
            'sections': [
                ('Proposed Architecture', [
                    'We propose a multi-task transformer architecture that combines shared encoder layers with task-specific decoder heads. The shared encoder consists of 12 transformer layers with 768 hidden dimensions and 12 attention heads, pre-trained on a combination of general and biomedical text corpora.',
                    'Each task-specific head consists of a two-layer feed-forward network with GELU activation and dropout regularization. The architecture supports dynamic task weighting through a learned temperature parameter that adjusts the contribution of each task during training.',
                ]),
                ('Data Collection and Preprocessing', [
                    'Our training data comprises three sources: (1) the PubMed abstracts dataset containing 14.2 million biomedical abstracts, (2) the MIMIC-III clinical notes dataset with 2.1 million de-identified clinical notes, and (3) the BioASQ question-answering dataset with 3,742 annotated question-answer pairs.',
                    'Text preprocessing involved tokenization using a domain-specific SentencePiece model trained on biomedical text, with a vocabulary size of 32,000 tokens. We applied lowercasing, removal of non-ASCII characters, and sentence boundary detection using scispaCy.',
                ]),
                ('Training Procedure', [
                    'Training was conducted on 8 NVIDIA A100 GPUs using mixed-precision (FP16) training with gradient accumulation. We used the AdamW optimizer with a learning rate of 2e-5, warmup over 10% of total steps, and linear decay. The batch size was 32 per GPU with gradient accumulation steps of 4.',
                ]),
                ('Evaluation Metrics', [
                    'We evaluate our models using standard metrics: F1 score for named entity recognition, accuracy for text classification, exact match (EM) and F1 for question answering, and BLEU-4 for text generation tasks. Statistical significance is assessed using bootstrap resampling with 10,000 iterations.',
                ]),
            ],
            'figures': [
                ('Detailed architecture diagram of the proposed multi-task transformer model with shared encoder and task-specific heads', 4),
                ('Data preprocessing pipeline showing tokenization, filtering, and augmentation steps', 5),
            ],
        },
        {
            'title': 'Experimental Setup',
            'sections': [
                ('Datasets', [
                    'We evaluate our approach on six benchmark datasets spanning four NLU tasks. For named entity recognition: BC5CDR (Li et al., 2016) and NCBI Disease (Dogan et al., 2014). For relation extraction: ChemProt (Kringelum et al., 2016). For question answering: BioASQ Task 7b. For text classification: HoC (Baker et al., 2016) and LitCovid.',
                ]),
                ('Baseline Models', [
                    'We compare against five baseline models: (1) BioBERT v1.1 (Lee et al., 2020), (2) PubMedBERT (Gu et al., 2021), (3) SciBERT (Beltagy et al., 2019), (4) ClinicalBERT (Alsentzer et al., 2019), and (5) a vanilla BERT-base model fine-tuned on each task independently.',
                ]),
                ('Hyperparameter Search', [
                    'We performed hyperparameter optimization using Optuna with Tree-structured Parzen Estimator (TPE) sampling. The search space included learning rate [1e-5, 5e-5], batch size {16, 32, 64}, dropout rate [0.1, 0.3], and number of fine-tuning epochs {3, 5, 10}. Each configuration was evaluated using 3-fold cross-validation.',
                ]),
                ('Implementation Details', [
                    'All experiments were implemented using PyTorch 2.0 with the HuggingFace Transformers library version 4.28. Model checkpoints were saved every 500 steps, and early stopping was applied based on validation loss with a patience of 5 epochs. Gradient clipping was set to a maximum norm of 1.0 to prevent training instabilities.',
                ]),
            ],
            'figures': [
                ('Distribution of dataset sizes across the six benchmark datasets used in evaluation', 6),
            ],
        },
        {
            'title': 'Results and Analysis',
            'sections': [
                ('Main Results', [
                    'Table 5.1 presents the main results across all six benchmark datasets. Our multi-task model achieves state-of-the-art performance on four of six datasets, with statistically significant improvements (p < 0.05) on BC5CDR (F1: 89.7 vs. 88.2), ChemProt (F1: 77.3 vs. 75.8), and BioASQ (F1: 51.2 vs. 48.9).',
                    'On the remaining two datasets, our model performs within 0.3 points of the best baseline, suggesting that multi-task learning provides consistent benefits without significant trade-offs on any individual task.',
                ]),
                ('Ablation Study', [
                    'We conducted an extensive ablation study to understand the contribution of each component. Removing the dynamic task weighting mechanism reduced average F1 by 1.8 points. Replacing the domain-specific vocabulary with a general vocabulary decreased performance by 2.1 points on biomedical NER tasks.',
                ]),
                ('Error Analysis', [
                    'Qualitative error analysis reveals three primary failure modes: (1) ambiguous entity boundaries in nested entities (accounting for 34% of NER errors), (2) implicit relations requiring world knowledge (41% of RE errors), and (3) multi-hop reasoning questions (28% of QA errors).',
                ]),
                ('Computational Efficiency', [
                    'Despite the multi-task overhead, our model achieves a 2.3x speedup in total training time compared to training separate models for each task. Inference latency increases by only 8% due to the lightweight task-specific heads, while GPU memory usage decreases by 40% compared to maintaining separate model instances.',
                ]),
            ],
            'figures': [
                ('Performance comparison bar chart across all six datasets for each model', 7),
                ('Learning curves showing validation loss over training steps for single-task versus multi-task configurations', 8),
            ],
        },
        {
            'title': 'Discussion',
            'sections': [
                ('Implications for Biomedical NLP', [
                    'Our results demonstrate that multi-task learning provides a practical and effective approach for biomedical NLU. The shared encoder learns representations that capture domain-specific linguistic patterns while maintaining generalizability across task types. This finding has important implications for resource-constrained settings where labeled data is scarce.',
                ]),
                ('Limitations', [
                    'Several limitations should be noted. First, our evaluation is limited to English-language biomedical text; the generalizability to other languages remains untested. Second, the computational requirements for pre-training remain substantial, potentially limiting accessibility for smaller research groups. Third, our domain adaptation approach may not transfer effectively to highly specialized subdomains with unique terminology.',
                ]),
                ('Comparison with Recent Work', [
                    'Concurrent work by Zhang et al. (2024) proposes a similar multi-task framework but focuses on clinical text. Their results on MIMIC-III tasks complement our findings, suggesting that the multi-task approach generalizes across biomedical subdomains. However, their architecture lacks the dynamic task weighting mechanism that we find crucial for balancing heterogeneous task objectives.',
                ]),
            ],
            'figures': [
                ('Attention visualization heatmap showing how the model attends to biomedical entities in context', 9),
            ],
        },
        {
            'title': 'Conclusion and Future Work',
            'sections': [
                ('Summary of Contributions', [
                    'This thesis makes three primary contributions: (1) a multi-task transformer architecture with dynamic task weighting for biomedical NLU, achieving state-of-the-art results on four benchmark datasets; (2) a comprehensive analysis of multi-task learning dynamics in domain-specific settings; and (3) practical guidelines for applying transfer learning to specialized text domains.',
                ]),
                ('Future Directions', [
                    'Future work should explore several promising directions. First, extending the multi-task framework to multilingual biomedical NLU could address the critical need for NLP tools in non-English medical literature. Second, incorporating structured knowledge from biomedical ontologies (e.g., UMLS, Gene Ontology) could improve entity disambiguation and relation extraction.',
                    'Third, the application of instruction tuning and in-context learning techniques from large language models to our multi-task framework presents an exciting avenue for few-shot biomedical NLU. Finally, deploying and evaluating the model in real clinical workflows would provide valuable insights into practical utility and user acceptance.',
                ]),
                ('Broader Impact', [
                    'The broader impact of this work extends beyond the immediate technical contributions. By demonstrating the effectiveness of multi-task learning in specialized domains, we provide a framework that can be adapted to other high-stakes fields such as legal document analysis, financial text processing, and environmental monitoring. The reduced computational requirements of multi-task approaches also contribute to more sustainable AI research practices.',
                ]),
            ],
            'figures': [
                ('Roadmap diagram showing proposed future research directions and their interconnections', 10),
            ],
        },
    ]

    figure_counter = 0

    for ch_idx, chapter in enumerate(chapters, 1):
        # Chapter heading (Heading 1)
        doc.add_heading(f'Chapter {ch_idx}: {chapter["title"]}', level=1)

        for sec_title, sec_paragraphs in chapter['sections']:
            # Sub-section heading (Heading 2)
            doc.add_heading(f'{ch_idx}.{chapter["sections"].index((sec_title, sec_paragraphs)) + 1} {sec_title}', level=2)

            for para_text in sec_paragraphs:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)

        # Add figures for this chapter
        for fig_caption, fig_num in chapter.get('figures', []):
            figure_counter += 1

            # Add some space before figure
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(12)

            # Figure placeholder (a bordered paragraph simulating a figure)
            fig_para = doc.add_paragraph()
            fig_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            fig_para.paragraph_format.space_before = Pt(6)
            run = fig_para.add_run(f'[Figure {figure_counter} placeholder image]')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True

            # Figure caption with "Figure X:" prefix (this is what Table of Figures should list)
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption_para.paragraph_format.space_after = Pt(12)
            caption_para.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else doc.styles['Normal']
            run = caption_para.add_run(f'Figure {figure_counter}: {fig_caption}')
            run.font.size = Pt(10)
            run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
