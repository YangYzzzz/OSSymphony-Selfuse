"""
Initial Setup: Offer letter with 'Sarah Johnson' in regular weight
Task ID: writer_hr_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Company Header ---
    header_para = doc.add_paragraph()
    header_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.paragraph_format.space_after = Pt(4)
    run = header_para.add_run("Meridian Technologies Inc.")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    addr_para = doc.add_paragraph()
    addr_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr_para.paragraph_format.space_after = Pt(2)
    run = addr_para.add_run("2400 Innovation Boulevard, Suite 500")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    addr_para2 = doc.add_paragraph()
    addr_para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr_para2.paragraph_format.space_after = Pt(12)
    run = addr_para2.add_run("Austin, TX 78701 | hr@meridiantech.com | (512) 555-0198")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.add_run("March 18, 2025")
    run.font.size = Pt(11)

    # --- Title --- (Occurrence 1: "Sarah Johnson" in the title line)
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after = Pt(12)
    run = title_para.add_run("OFFER OF EMPLOYMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # --- Greeting (occurrence 1) ---
    greeting = doc.add_paragraph()
    greeting.paragraph_format.space_after = Pt(6)
    run1 = greeting.add_run("Dear ")
    run1.font.size = Pt(11)
    run2 = greeting.add_run("Sarah Johnson")  # Occurrence 1
    run2.font.size = Pt(11)
    run2.bold = False
    run3 = greeting.add_run(",")
    run3.font.size = Pt(11)

    # --- Opening paragraph (occurrence 2) ---
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run("We are pleased to extend this formal offer of employment to ")
    r1.font.size = Pt(11)
    r2 = p1.add_run("Sarah Johnson")  # Occurrence 2
    r2.font.size = Pt(11)
    r2.bold = False
    r3 = p1.add_run(
        " for the position of Senior Software Engineer in our Cloud Infrastructure "
        "division. After a thorough review of your qualifications and interview "
        "performance, we are confident that your expertise in distributed systems "
        "and microservices architecture will be a tremendous asset to our team."
    )
    r3.font.size = Pt(11)

    # --- Position Details heading ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    run = h1.add_run("Position Details")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # --- Position details ---
    details = [
        ("Title:", "Senior Software Engineer"),
        ("Department:", "Cloud Infrastructure"),
        ("Reports To:", "David Kim, VP of Engineering"),
        ("Start Date:", "April 14, 2025"),
        ("Location:", "Austin, TX (Hybrid \u2013 3 days on-site)"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rv = p.add_run(value)
        rv.font.size = Pt(11)

    # --- Compensation heading ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    run = h2.add_run("Compensation & Benefits")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    comp_items = [
        "Base Salary: $145,000 per annum, paid bi-weekly",
        "Annual Bonus: Up to 15% of base salary, based on performance metrics",
        "Equity: 8,000 stock options vesting over 4 years with a 1-year cliff",
        "Health Insurance: Comprehensive medical, dental, and vision coverage (employer pays 90%)",
        "401(k): Company match of 50% up to 6% of salary",
        "PTO: 20 days paid vacation plus 10 company holidays",
        "Professional Development: $3,000 annual learning stipend",
    ]
    for item in comp_items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    # --- Conditions heading ---
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    run = h3.add_run("Conditions of Employment")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    cond_para = doc.add_paragraph()
    cond_para.paragraph_format.space_after = Pt(6)
    run = cond_para.add_run(
        "This offer is contingent upon satisfactory completion of a background check, "
        "verification of employment eligibility, and execution of our standard "
        "Confidentiality and Intellectual Property Agreement. Employment at Meridian "
        "Technologies is at-will, meaning either party may terminate the relationship "
        "at any time with or without cause."
    )
    run.font.size = Pt(11)

    # --- Acceptance paragraph (occurrence 3) ---
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    run = h4.add_run("Acceptance")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    accept_para = doc.add_paragraph()
    accept_para.paragraph_format.space_after = Pt(6)
    r1 = accept_para.add_run(
        "To accept this offer, please sign and return this letter by April 4, 2025. "
        "Should you have any questions regarding the terms outlined above, please do "
        "not hesitate to reach out to our HR department. We are eager to welcome "
    )
    r1.font.size = Pt(11)
    r2 = accept_para.add_run("Sarah Johnson")  # Occurrence 3
    r2.font.size = Pt(11)
    r2.bold = False
    r3 = accept_para.add_run(" to the Meridian Technologies family.")
    r3.font.size = Pt(11)

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(18)
    closing.paragraph_format.space_after = Pt(2)
    run = closing.add_run("Sincerely,")
    run.font.size = Pt(11)

    sig = doc.add_paragraph()
    sig.paragraph_format.space_after = Pt(2)
    run = sig.add_run("Elena Martinez")
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(18)
    run = sig_title.add_run("Director of Human Resources\nMeridian Technologies Inc.")
    run.font.size = Pt(11)

    # --- Signature line (occurrence 4) ---
    sig_line = doc.add_paragraph()
    sig_line.paragraph_format.space_before = Pt(24)
    sig_line.paragraph_format.space_after = Pt(2)
    r1 = sig_line.add_run("Accepted by: ")
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = sig_line.add_run("Sarah Johnson")  # Occurrence 4
    r2.font.size = Pt(11)
    r2.bold = False

    date_line = doc.add_paragraph()
    date_line.paragraph_format.space_after = Pt(2)
    r1 = date_line.add_run("Date: _______________")
    r1.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
