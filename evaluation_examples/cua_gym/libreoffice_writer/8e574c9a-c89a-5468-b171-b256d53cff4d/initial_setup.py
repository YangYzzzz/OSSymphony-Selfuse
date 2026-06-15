"""
Initial Setup: Company announcement document (no image inserted yet).
Task ID: osworld_writer_image_insertion_003
Domain: libreoffice_writer

Creates:
  - /home/user/osworld_writer_image_insertion_003.docx  — one-page company announcement
  - /home/user/banner.jpg  — banner image to be inserted by the agent
"""

import os
import shlex
import subprocess
import time
import struct
import zlib

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_image_insertion_003'
OUTPUT  = f'{WORKDIR}/{TASK_ID}.docx'
BANNER  = f'{WORKDIR}/banner.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_banner_jpg(path: str):
    """Create a small, valid JPEG banner image using raw bytes (no external deps on VM)."""
    # Minimal valid JPEG (1x1 blue pixel) — enough for python-docx to embed
    # We use a small solid-color JPEG created via struct/zlib-free approach:
    # Standard minimal JPEG bytes for a 100x30 blue rectangle
    import io
    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (400, 100), color=(30, 90, 180))
        draw = ImageDraw.Draw(img)
        draw.rectangle([0, 0, 399, 99], outline=(255, 255, 255), width=3)
        draw.text((20, 35), "Nexara Technologies", fill=(255, 255, 255))
        img.save(path, 'JPEG', quality=85)
    except Exception:
        # Fallback: write a minimal valid JPEG (white 10x10)
        jpeg_bytes = (
            b'\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00'
            b'\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t'
            b'\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a'
            b'\x1f\x1e\x1d\x1a\x1c\x1c $.\' ",#\x1c\x1c(7),01444\x1f\'9=82<.342\x1e\x1e'
            b'\xff\xc0\x00\x0b\x08\x00\x0a\x00\x0a\x01\x01\x11\x00'
            b'\xff\xc4\x00\x1f\x00\x00\x01\x05\x01\x01\x01\x01\x01\x01\x00\x00'
            b'\x00\x00\x00\x00\x00\x00\x01\x02\x03\x04\x05\x06\x07\x08\t\n\x0b'
            b'\xff\xc4\x00\xb5\x10\x00\x02\x01\x03\x03\x02\x04\x03\x05\x05\x04'
            b'\x04\x00\x00\x01}\x01\x02\x03\x00\x04\x11\x05\x12!1A\x06\x13Qa'
            b'\x07"q\x142\x81\x91\xa1\x08#B\xb1\xc1\x15R\xd1\xf0$3br'
            b'\x82\t\n\x16\x17\x18\x19\x1a%&\'()*456789:CDEFGHIJSTUVWXYZ'
            b'cdefghijstuvwxyz\x83\x84\x85\x86\x87\x88\x89\x8a\x92\x93\x94'
            b'\x95\x96\x97\x98\x99\x9a\xa2\xa3\xa4\xa5\xa6\xa7\xa8\xa9\xaa'
            b'\xb2\xb3\xb4\xb5\xb6\xb7\xb8\xb9\xba\xc2\xc3\xc4\xc5\xc6\xc7'
            b'\xc8\xc9\xca\xd2\xd3\xd4\xd5\xd6\xd7\xd8\xd9\xda\xe1\xe2\xe3'
            b'\xe4\xe5\xe6\xe7\xe8\xe9\xea\xf1\xf2\xf3\xf4\xf5\xf6\xf7\xf8'
            b'\xf9\xfa\xff\xda\x00\x08\x01\x01\x00\x00?\x00\xfb\xd3P\x00\x00'
            b'\x00\x1f\xff\xd9'
        )
        with open(path, 'wb') as f:
            f.write(jpeg_bytes)
    print(f'Banner image created: {path}')


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page setup: letter size with standard margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Heading ---
    heading = doc.add_heading('Company Announcement', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Date / sub-heading ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('March 2025 | Internal Communication')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()  # spacer

    # --- Body paragraphs ---
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'We are thrilled to announce that Nexara Technologies has achieved a significant milestone '
        'in our ongoing commitment to innovation and excellence. Effective April 1, 2025, our '
        'organization will be launching a comprehensive digital transformation initiative across '
        'all regional offices.'
    )
    intro_run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    p2 = doc.add_paragraph()
    p2.add_run('Leadership Update').bold = True
    p3 = doc.add_paragraph()
    p3.add_run(
        'We welcome Dr. Elena Vasquez as our new Chief Technology Officer. Dr. Vasquez brings '
        'over 18 years of experience leading technology teams at Fortune 500 companies. She will '
        'be overseeing the technology roadmap and ensuring our systems remain world-class. '
        'Dr. Vasquez will officially join on April 15, 2025.'
    ).font.size = Pt(11)

    doc.add_paragraph()  # spacer

    p4 = doc.add_paragraph()
    p4.add_run('Strategic Partnerships').bold = True
    p5 = doc.add_paragraph()
    p5.add_run(
        'Nexara Technologies has signed a strategic partnership agreement with CloudNova Solutions '
        'to enhance our cloud infrastructure. This partnership is expected to reduce operational '
        'costs by 23% while improving system reliability and scalability for our 4,200 global '
        'customers. The partnership officially commences on May 1, 2025.'
    ).font.size = Pt(11)

    doc.add_paragraph()  # spacer

    p6 = doc.add_paragraph()
    p6.add_run('Employee Benefits').bold = True
    p7 = doc.add_paragraph()
    p7.add_run(
        'As part of our People-First strategy, Human Resources will be rolling out an enhanced '
        'benefits package starting Q2 2025. Highlights include a 10% increase in professional '
        'development budgets, flexible remote work options for eligible positions, and expanded '
        'healthcare coverage for employees and their families. Full details will be shared by '
        'HR Director Marcus Webb on March 28, 2025.'
    ).font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cr = closing.add_run('— Nexara Technologies Leadership Team —')
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.save(OUTPUT)
    print(f'Initial document created: {OUTPUT}')

    # Create banner.jpg in home directory (not inserted in doc yet)
    create_banner_jpg(BANNER)

    # GUI-ready startup: open document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
