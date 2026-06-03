"""
Reward Script: Risk Register Document in LibreOffice Writer
Task ID: writer_wf_092
Domain: libreoffice_writer
Scoring:
  Component 1: Title present (0.15)
  Component 2: Project info table with 3 fields (0.15)
  Component 3: Risk register table header with 9 columns (0.20)
  Component 4: Risk register has 8 risk rows across 4 categories (0.20)
  Component 5: Risk matrix 5x5 grid (0.15)
  Component 6: Review log table (0.15)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_092'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have content (not blank)
    all_text = ' '.join(p.text for p in doc.paragraphs)
    if not all_text.strip() and len(doc.tables) == 0:
        print("FAIL: Document is blank — no text and no tables")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title present — "Project Risk Register - Smart City Initiative" (0.15 points)
    try:
        title_found = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            if 'project risk register' in text_lower and 'smart city' in text_lower:
                title_found = True
                break
        # Also check table cells in case title is in a table
        if not title_found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cl = cell.text.strip().lower()
                        if 'project risk register' in cl and 'smart city' in cl:
                            title_found = True
                            break
                    if title_found:
                        break
                if title_found:
                    break

        if title_found:
            print(f"PASS: Component 1 — Title found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title 'Project Risk Register - Smart City Initiative' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Project info table with Project Manager, Last Updated, Review Frequency (0.15 points)
    try:
        project_info_found = False
        required_fields = {'project manager', 'last updated', 'review frequency'}
        for table in doc.tables:
            found_fields = set()
            for row in table.rows:
                for cell in row.cells:
                    cl = cell.text.strip().lower()
                    for field in required_fields:
                        if field in cl:
                            found_fields.add(field)
            if found_fields == required_fields:
                project_info_found = True
                break

        if project_info_found:
            print(f"PASS: Component 2 — Project info table with all 3 fields (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Project info table missing or incomplete. Need: {required_fields}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Risk register table with 9-column header (0.20 points)
    try:
        risk_table = None
        expected_headers = ['risk id', 'category', 'description', 'likelihood', 'impact',
                            'risk score', 'mitigation', 'owner', 'status']
        for table in doc.tables:
            if len(table.columns) >= 9:
                # Check header row
                header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
                matched = 0
                for eh in expected_headers:
                    if any(eh in ht for ht in header_texts):
                        matched += 1
                if matched >= 7:  # at least 7 of 9 headers match
                    risk_table = table
                    break

        if risk_table is not None:
            print(f"PASS: Component 3 — Risk register table found with 9 columns and correct headers (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — No risk register table with 9 columns and matching headers found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Risk register has 8 risk entries across 4 categories (Technical, Financial, Operational, Legal) (0.20 points)
    try:
        if risk_table is not None:
            data_rows = list(risk_table.rows)[1:]  # skip header
            num_data_rows = len(data_rows)

            # Check categories
            required_categories = {'technical', 'financial', 'operational', 'legal'}
            found_categories = set()
            for row in data_rows:
                for cell in row.cells:
                    cl = cell.text.strip().lower()
                    for cat in required_categories:
                        if cat in cl:
                            found_categories.add(cat)

            score_4 = 0.0
            # At least 8 data rows
            if num_data_rows >= 8:
                score_4 += 0.10
                print(f"  Component 4a: {num_data_rows} risk rows (>= 8) — 0.10 pts")
            else:
                print(f"  Component 4a: Only {num_data_rows} risk rows (need >= 8)")

            # All 4 categories present
            if found_categories == required_categories:
                score_4 += 0.10
                print(f"  Component 4b: All 4 categories found: {found_categories} — 0.10 pts")
            else:
                print(f"  Component 4b: Missing categories. Found: {found_categories}, Need: {required_categories}")

            if score_4 > 0:
                print(f"PASS: Component 4 — Risk data ({score_4} pts)")
                total_score += score_4
            else:
                print(f"FAIL: Component 4 — Risk data insufficient")
        else:
            print(f"FAIL: Component 4 — No risk register table found (depends on Component 3)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Risk Matrix — a 5x5 grid showing Likelihood vs Impact (0.15 points)
    try:
        matrix_found = False
        for table in doc.tables:
            # Skip the risk register table
            if risk_table is not None and table._element is risk_table._element:
                continue
            # A risk matrix should have rows >= 6 (header + 5 likelihood levels)
            # and cols >= 6 (row labels + 5 impact levels)
            if len(table.rows) >= 6 and len(table.columns) >= 6:
                # Check for matrix-like content: numbers that are products of likelihood x impact
                all_text = ''
                for row in table.rows:
                    for cell in row.cells:
                        all_text += ' ' + cell.text.strip().lower()
                # Look for markers of a risk matrix
                has_likelihood = any(kw in all_text for kw in ['likelihood', 'likely', 'rare', 'possible', 'certain', 'unlikely'])
                has_impact = any(kw in all_text for kw in ['impact', 'negligible', 'minor', 'moderate', 'major', 'catastrophic'])
                # Also check for numeric grid pattern (1-25 range)
                has_numbers = False
                num_count = 0
                for row in table.rows:
                    for cell in row.cells:
                        try:
                            v = int(cell.text.strip())
                            if 1 <= v <= 25:
                                num_count += 1
                        except ValueError:
                            pass
                if num_count >= 15:  # at least 15 numeric cells (5x5 grid = 25)
                    has_numbers = True

                if (has_likelihood or has_impact) and has_numbers:
                    matrix_found = True
                    break

        if matrix_found:
            print(f"PASS: Component 5 — Risk matrix (5x5 grid) found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Risk matrix not found or incomplete")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Review log table with Date, Reviewer, Changes, Next Review columns (0.15 points)
    try:
        review_log_found = False
        review_headers = ['date', 'reviewer', 'changes', 'next review']
        for table in doc.tables:
            # Skip risk register and matrix tables
            if risk_table is not None and table._element is risk_table._element:
                continue
            # Check header row for review log columns
            if len(table.rows) >= 2:
                header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
                header_combined = ' '.join(header_texts)
                matched = sum(1 for rh in review_headers if rh in header_combined)
                if matched >= 3:  # at least 3 of 4 review log headers
                    # Also verify it has data rows (not just header)
                    if len(table.rows) >= 2:
                        review_log_found = True
                        break

        if review_log_found:
            print(f"PASS: Component 6 — Review log table found with correct columns (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 6 — Review log table not found or missing columns")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
