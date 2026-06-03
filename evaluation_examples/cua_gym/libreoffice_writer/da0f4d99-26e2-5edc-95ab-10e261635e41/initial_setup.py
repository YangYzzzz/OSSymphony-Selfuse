"""
Initial Setup: Technical manual with Figure 1 caption and preceding paragraph containing
'The diagram below illustrates the process flow.'
Task ID: osworld_writer_bibliography_crossref_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_figure_caption(doc, caption_text, figure_num):
    """Add a figure caption paragraph with SEQ field for figure numbering."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)

    # Add "Figure " prefix run
    run_prefix = para.add_run("Figure ")
    run_prefix.bold = True
    run_prefix.font.size = Pt(10)

    # Add SEQ field for auto-numbering (Figure 1, Figure 2, etc.)
    # This is an fldChar-based SEQ field
    run_seq = para.add_run()
    run_seq.bold = True
    run_seq.font.size = Pt(10)

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_seq._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' SEQ Figure \\* ARABIC '
    run_seq._r.append(instrText)

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run_seq._r.append(fldChar_separate)

    # Cached value for the field
    run_cache = para.add_run(str(figure_num))
    run_cache.bold = True
    run_cache.font.size = Pt(10)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    # Add end fldChar to a new run
    run_end = para.add_run()
    run_end.bold = True
    run_end.font.size = Pt(10)
    run_end._r.append(fldChar_end)

    # Add caption text after the figure number
    run_colon = para.add_run(f": {caption_text}")
    run_colon.bold = True
    run_colon.font.size = Pt(10)

    return para


def create_initial():
    doc = Document()

    # --- Document Title ---
    title_para = doc.add_heading("Technical Manual: Data Processing System", level=0)
    title_para.paragraph_format.space_after = Pt(12)

    # --- Chapter 1 ---
    doc.add_heading("Chapter 1: System Overview", level=1)

    intro_para = doc.add_paragraph(
        "This manual provides a comprehensive guide to the Data Processing System (DPS) "
        "version 3.2. The system is designed to handle high-volume data ingestion, "
        "transformation, and output across multiple data formats and protocols."
    )
    intro_para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "The DPS integrates seamlessly with existing enterprise infrastructure, supporting "
        "SQL and NoSQL databases, REST and SOAP APIs, and batch file processing pipelines. "
        "System administrators can configure workflows through the web-based management console."
    )

    # --- Section 1.1 ---
    doc.add_heading("1.1 Architecture Overview", level=2)

    arch_para = doc.add_paragraph(
        "The system architecture follows a modular microservices design, with each component "
        "responsible for a specific stage in the data pipeline. The core components include "
        "the Ingestion Layer, the Processing Engine, and the Output Handler."
    )
    arch_para.paragraph_format.space_after = Pt(6)

    # Paragraph BEFORE Figure 1 — contains the target sentence
    flow_intro = doc.add_paragraph(
        "The system processes data through a series of well-defined stages to ensure "
        "accuracy and consistency. The diagram below illustrates the process flow."
    )
    flow_intro.paragraph_format.space_after = Pt(6)

    # --- Placeholder image area (represented as a bordered paragraph) ---
    figure_placeholder = doc.add_paragraph()
    figure_placeholder.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    figure_placeholder.paragraph_format.space_before = Pt(12)
    figure_placeholder.paragraph_format.space_after = Pt(6)
    run_fig = figure_placeholder.add_run("[Process Flow Diagram]")
    run_fig.font.size = Pt(12)
    run_fig.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run_fig.italic = True

    # Figure 1 Caption
    add_figure_caption(doc, "Process Flow Diagram", 1)

    # --- Section 1.2 ---
    doc.add_heading("1.2 Data Ingestion Layer", level=2)

    doc.add_paragraph(
        "The Ingestion Layer accepts data from multiple sources including relational databases, "
        "flat files (CSV, JSON, XML), and streaming interfaces. Each data source is configured "
        "through the Source Configuration Manager."
    )

    doc.add_paragraph(
        "Supported database connectors include: Oracle 19c+, Microsoft SQL Server 2019+, "
        "PostgreSQL 13+, MySQL 8+, and MongoDB 5+. Custom connectors can be developed using "
        "the provided SDK and connector API."
    )

    # --- Section 1.3 ---
    doc.add_heading("1.3 Processing Engine", level=2)

    doc.add_paragraph(
        "The Processing Engine applies transformation rules defined in the workflow configuration. "
        "Rules are evaluated in sequence and support conditional branching, data enrichment "
        "from lookup tables, and complex aggregation functions."
    )

    doc.add_paragraph(
        "Performance benchmarks indicate the engine can process up to 50,000 records per second "
        "on standard hardware (8-core CPU, 32 GB RAM). Horizontal scaling is supported through "
        "the cluster configuration module."
    )

    # --- Chapter 2 ---
    doc.add_heading("Chapter 2: Installation and Configuration", level=1)

    doc.add_paragraph(
        "Before installing the Data Processing System, ensure that the target server meets "
        "the minimum hardware and software requirements specified in Appendix A."
    )

    doc.add_paragraph(
        "The installation package includes the core application server, the management console, "
        "monitoring agents, and sample workflow configurations. Installation typically takes "
        "15-30 minutes depending on system performance."
    )

    # --- Table of system requirements ---
    doc.add_heading("2.1 System Requirements", level=2)

    req_table = doc.add_table(rows=5, cols=2)
    req_table.style = "Table Grid"

    headers = [("Component", "Minimum Requirement")]
    rows_data = [
        ("CPU", "8-core x86_64 processor"),
        ("RAM", "32 GB DDR4"),
        ("Storage", "500 GB SSD"),
        ("OS", "Ubuntu 20.04 LTS or RHEL 8+"),
    ]

    # Header row
    header_row = req_table.rows[0]
    for i, h in enumerate(headers[0]):
        cell = header_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(11)

    # Data rows
    for row_idx, (col1, col2) in enumerate(rows_data, 1):
        req_table.rows[row_idx].cells[0].text = col1
        req_table.rows[row_idx].cells[1].text = col2

    doc.add_paragraph()  # Spacing after table

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
