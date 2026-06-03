"""
Initial Setup: Create a Writer document with a code snippet in Liberation Sans 12pt.
Task ID: writer_tech_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Setting Up a Python Development Environment', level=1)

    # --- Introductory paragraph ---
    intro = doc.add_paragraph()
    run = intro.add_run(
        'Python is one of the most popular programming languages for web development, '
        'data science, and automation. This guide walks you through installing Python 3 '
        'on a Debian-based Linux system and configuring your development environment.'
    )
    run.font.name = 'Liberation Sans'
    run.font.size = Pt(12)

    # --- Paragraph with the code snippet ---
    p2 = doc.add_paragraph()
    run1 = p2.add_run('To install Python 3, open a terminal and run the command ')
    run1.font.name = 'Liberation Sans'
    run1.font.size = Pt(12)

    # The code snippet itself - in Liberation Sans 12pt (initial state)
    code_run = p2.add_run('sudo apt install python3')
    code_run.font.name = 'Liberation Sans'
    code_run.font.size = Pt(12)

    run2 = p2.add_run('. This will download and install the latest Python 3 package '
                       'available in your distribution\'s repository.')
    run2.font.name = 'Liberation Sans'
    run2.font.size = Pt(12)

    # --- Additional paragraphs for realism ---
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        'After the installation completes, verify it by checking the version number. '
        'You should see output similar to Python 3.11.2 or later depending on your '
        'distribution. It is recommended to also install pip, the Python package '
        'manager, which lets you install third-party libraries from the Python Package Index.'
    )
    run3.font.name = 'Liberation Sans'
    run3.font.size = Pt(12)

    p4 = doc.add_paragraph()
    run4 = p4.add_run(
        'Next, consider setting up a virtual environment for each project. Virtual '
        'environments keep dependencies isolated so that different projects can use '
        'different library versions without conflicts. The venv module, included with '
        'Python 3.3 and later, is the standard tool for creating lightweight virtual '
        'environments.'
    )
    run4.font.name = 'Liberation Sans'
    run4.font.size = Pt(12)

    p5 = doc.add_paragraph()
    run5 = p5.add_run(
        'For teams working on larger projects, you may also want to explore tools like '
        'Poetry or Pipenv, which combine dependency management with virtual environment '
        'creation. These tools generate lock files that ensure reproducible builds across '
        'all developer machines and CI pipelines.'
    )
    run5.font.name = 'Liberation Sans'
    run5.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
