"""
Initial Setup: Social Media Report - Plain Text Version (Pre-Task State)
Task ID: writer_mktg_012
Domain: libreoffice_writer

Creates a plain text .docx file with unstructured social media analytics data
on ~/Desktop/ that the agent needs to format into a polished report.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'social_media_report_feb'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Remove default styles - use plain paragraphs at 11pt
    def add_plain(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        return para

    # Header block - plain text, no formatting
    add_plain("February 2026 Social Media Performance")
    add_plain("Exported from Analytics Dashboard - February 1-28, 2026")
    add_plain("")

    # Instagram block
    add_plain("Instagram")
    add_plain("Followers 45,200 (+8.2%)")
    add_plain("Engagement Rate 4.7% (+0.3%)")
    add_plain("Posts 28")
    add_plain("Story Views 312,450")
    add_plain("")

    # Twitter/X block
    add_plain("Twitter/X")
    add_plain("Followers 18,750 (+2.1%)")
    add_plain("Engagement Rate 1.8% (-0.4%)")
    add_plain("Posts 62")
    add_plain("Impressions 485,200")
    add_plain("")

    # LinkedIn block
    add_plain("LinkedIn")
    add_plain("Followers 9,340 (+5.6%)")
    add_plain("Engagement Rate 3.2% (+0.8%)")
    add_plain("Posts 15")
    add_plain("Profile Views 4,210")
    add_plain("")

    # TikTok block
    add_plain("TikTok")
    add_plain("Followers 22,100 (+15.3%)")
    add_plain("Engagement Rate 8.4% (+1.2%)")
    add_plain("Posts 19")
    add_plain("Video Views 1,245,000")
    add_plain("")

    # Facebook block
    add_plain("Facebook")
    add_plain("Followers 31,500 (-0.6%)")
    add_plain("Engagement Rate 0.9% (-0.2%)")
    add_plain("Posts 24")
    add_plain("Reach 198,750")
    add_plain("")

    # Summary metrics
    add_plain("Summary")
    add_plain("Total Reach across all platforms: 2,241,400")
    add_plain("Total Impressions across all platforms: 3,876,500")
    add_plain("Combined Follower Count: 126,890")
    add_plain("Average Engagement Rate: 3.8%")

    # Ensure output directory exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
