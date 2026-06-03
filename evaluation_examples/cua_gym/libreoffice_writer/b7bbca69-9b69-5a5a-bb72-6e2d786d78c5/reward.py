"""
Reward Script: Social Media Analytics Report Formatting
Task ID: writer_mktg_012
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Title paragraph is 18pt bold and centered
  Component 2 (0.15): Subtitle paragraph is 12pt italic and centered
  Component 3 (0.25): Table with 5 data rows and correct column headers
  Component 4 (0.15): Header row bold with gray (#E0E0E0) background shading
  Component 5 (0.10): MoM Change column: positive changes green (#2E7D32), negative red (#C62828)
  Component 6 (0.10): Summary metrics (Total Reach, Total Impressions) are bold
"""

import os
import math

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_012'
FILE_NAME = 'social_media_report_feb.docx'

# Target colors from task context
GREEN_COLOR = RGBColor(0x2E, 0x7D, 0x32)   # #2E7D32 — positive MoM
RED_COLOR   = RGBColor(0xC6, 0x28, 0x28)   # #C62828 — negative MoM
GRAY_FILL   = 'E0E0E0'                      # gray header background

EXPECTED_HEADERS = ['Platform', 'Followers', 'MoM Change', 'Engagement Rate', 'Posts']
EXPECTED_PLATFORMS = {'Instagram', 'Twitter/X', 'LinkedIn', 'TikTok', 'Facebook'}


def color_distance(c1: RGBColor, r: int, g: int, b: int) -> float:
    """Euclidean RGB distance between a RGBColor and an (r,g,b) tuple.
    RGBColor is a tuple subclass: c1[0]=red, c1[1]=green, c1[2]=blue."""
    return math.sqrt((c1[0] - r)**2 + (c1[1] - g)**2 + (c1[2] - b)**2)


def get_cell_fill(cell) -> str:
    """Return the fill hex string for a table cell, or empty string if none."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return ''
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return ''
    fill_val = shd.get(qn('w:fill'))
    return fill_val.upper() if fill_val else ''


def verify_task(file_path: str) -> float:
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Title paragraph (0.25 pts)
    # Must have text matching "Social Media Performance Report — February 2026",
    # be 18pt bold, and centered alignment.
    # In initial state: title is "February 2026 Social Media Performance", 11pt,
    # not bold, not centered — so this WILL fail on initial.
    # -----------------------------------------------------------------------
    try:
        title_found = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if 'Social Media Performance Report' in text and 'February 2026' in text:
                # Check centered alignment
                alignment_ok = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                # Check 18pt bold across runs
                all_runs = [r for r in para.runs if r.text.strip()]
                if not all_runs:
                    break
                size_ok = all(
                    r.font.size is not None and abs(r.font.size.pt - 18.0) < 0.5
                    for r in all_runs
                )
                bold_ok = all(r.bold is True for r in all_runs)
                if alignment_ok and size_ok and bold_ok:
                    title_found = True
                    print(f"PASS: Component 1 — Title is 18pt bold centered (0.25 pts)")
                    total_score += 0.25
                else:
                    issues = []
                    if not alignment_ok:
                        issues.append(f"alignment={para.paragraph_format.alignment} (expected CENTER)")
                    if not size_ok:
                        sizes = [r.font.size.pt if r.font.size else None for r in all_runs]
                        issues.append(f"sizes={sizes} (expected 18.0)")
                    if not bold_ok:
                        bolds = [r.bold for r in all_runs]
                        issues.append(f"bold={bolds} (expected all True)")
                    print(f"FAIL: Component 1 — Title found but formatting wrong: {'; '.join(issues)}")
                break
        if not title_found and total_score < 0.25:
            # Check if the issue was title not found at all vs formatting
            titles_present = [para.text.strip() for para in doc.paragraphs if 'Social Media Performance Report' in para.text]
            if not titles_present:
                print(f"FAIL: Component 1 — Title 'Social Media Performance Report — February 2026' not found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Subtitle paragraph (0.15 pts)
    # Must have text "Marketing Department | Monthly Review",
    # be 12pt italic, and centered alignment.
    # In initial state: no such subtitle exists — fails on initial.
    # -----------------------------------------------------------------------
    try:
        subtitle_found = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if 'Marketing Department' in text and 'Monthly Review' in text:
                alignment_ok = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                all_runs = [r for r in para.runs if r.text.strip()]
                if not all_runs:
                    break
                size_ok = all(
                    r.font.size is not None and abs(r.font.size.pt - 12.0) < 0.5
                    for r in all_runs
                )
                italic_ok = all(r.italic is True for r in all_runs)
                if alignment_ok and size_ok and italic_ok:
                    subtitle_found = True
                    print(f"PASS: Component 2 — Subtitle is 12pt italic centered (0.15 pts)")
                    total_score += 0.15
                else:
                    issues = []
                    if not alignment_ok:
                        issues.append(f"alignment={para.paragraph_format.alignment} (expected CENTER)")
                    if not size_ok:
                        sizes = [r.font.size.pt if r.font.size else None for r in all_runs]
                        issues.append(f"sizes={sizes} (expected 12.0)")
                    if not italic_ok:
                        italics = [r.italic for r in all_runs]
                        issues.append(f"italic={italics} (expected all True)")
                    print(f"FAIL: Component 2 — Subtitle found but formatting wrong: {'; '.join(issues)}")
                break
        if not subtitle_found and total_score < 0.40:
            if not any('Marketing Department' in para.text for para in doc.paragraphs):
                print(f"FAIL: Component 2 — Subtitle 'Marketing Department | Monthly Review' not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Table with correct structure (0.25 pts)
    # Must have exactly 1 table, with columns Platform/Followers/MoM Change/
    # Engagement Rate/Posts, and 5 data rows for the 5 platforms.
    # In initial state: 0 tables — fails on initial.
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 3 — No table found in document (expected 1 table with 5 data rows)")
        else:
            table = doc.tables[0]
            # Check header row
            if len(table.rows) == 0:
                print(f"FAIL: Component 3 — Table has no rows")
            else:
                header_row = table.rows[0]
                header_texts = [cell.text.strip() for cell in header_row.cells]
                headers_ok = header_texts == EXPECTED_HEADERS

                # Check 5 data rows have the correct platform names
                data_rows = table.rows[1:]
                found_platforms = set()
                for row in data_rows:
                    platform_cell = row.cells[0].text.strip()
                    if platform_cell:
                        found_platforms.add(platform_cell)

                platforms_ok = found_platforms == EXPECTED_PLATFORMS
                row_count_ok = len(data_rows) == 5

                if headers_ok and platforms_ok and row_count_ok:
                    print(f"PASS: Component 3 — Table with 5 data rows and correct headers: {header_texts} (0.25 pts)")
                    total_score += 0.25
                else:
                    issues = []
                    if not headers_ok:
                        issues.append(f"headers={header_texts} (expected {EXPECTED_HEADERS})")
                    if not row_count_ok:
                        issues.append(f"data rows={len(data_rows)} (expected 5)")
                    if not platforms_ok:
                        missing = EXPECTED_PLATFORMS - found_platforms
                        extra = found_platforms - EXPECTED_PLATFORMS
                        if missing:
                            issues.append(f"missing platforms={missing}")
                        if extra:
                            issues.append(f"unexpected platforms={extra}")
                    print(f"FAIL: Component 3 — Table structure incorrect: {'; '.join(issues)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Header row bold with gray (#E0E0E0) background (0.15 pts)
    # In initial state: no table exists — fails on initial.
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 4 — No table found, cannot check header formatting")
        else:
            table = doc.tables[0]
            if len(table.rows) == 0:
                print(f"FAIL: Component 4 — Table has no rows")
            else:
                header_row = table.rows[0]
                all_bold = True
                all_gray = True
                bold_details = []
                fill_details = []
                for cell in header_row.cells:
                    # Check bold
                    cell_runs = [r for p in cell.paragraphs for r in p.runs if r.text.strip()]
                    if cell_runs:
                        cell_bold = all(r.bold is True for r in cell_runs)
                    else:
                        cell_bold = False
                    bold_details.append(cell_bold)
                    if not cell_bold:
                        all_bold = False

                    # Check gray fill
                    fill_hex = get_cell_fill(cell)
                    fill_details.append(fill_hex)
                    if fill_hex.upper() != GRAY_FILL.upper():
                        all_gray = False

                if all_bold and all_gray:
                    print(f"PASS: Component 4 — Header row is bold with gray (#E0E0E0) background (0.15 pts)")
                    total_score += 0.15
                else:
                    issues = []
                    if not all_bold:
                        issues.append(f"bold per cell={bold_details} (expected all True)")
                    if not all_gray:
                        issues.append(f"fill per cell={fill_details} (expected all E0E0E0)")
                    print(f"FAIL: Component 4 — Header row formatting wrong: {'; '.join(issues)}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    # Component 5: MoM Change column color coding (0.10 pts)
    # Positive changes (starting with '+') must be green (#2E7D32).
    # Negative changes (starting with '-') must be red (#C62828).
    # In initial state: no table — fails on initial.
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 5 — No table found, cannot check MoM coloring")
        else:
            table = doc.tables[0]
            mom_col_idx = None
            if len(table.rows) > 0:
                header_row = table.rows[0]
                for j, cell in enumerate(header_row.cells):
                    if 'MoM' in cell.text or 'mom' in cell.text.lower():
                        mom_col_idx = j
                        break

            if mom_col_idx is None:
                print(f"FAIL: Component 5 — 'MoM Change' column not found in header row")
            else:
                data_rows = table.rows[1:]
                color_errors = []
                color_ok_count = 0
                for row_idx, row in enumerate(data_rows):
                    if mom_col_idx >= len(row.cells):
                        continue
                    cell = row.cells[mom_col_idx]
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    is_positive = cell_text.startswith('+')
                    is_negative = cell_text.startswith('-')

                    all_runs = [r for p in cell.paragraphs for r in p.runs if r.text.strip()]
                    if not all_runs:
                        continue

                    for run in all_runs:
                        try:
                            run_color = run.font.color.rgb
                        except Exception:
                            run_color = None

                        if is_positive:
                            # Expect green #2E7D32
                            if run_color is not None and color_distance(run_color, 0x2E, 0x7D, 0x32) < 20:
                                color_ok_count += 1
                            else:
                                color_errors.append(f"Row {row_idx+1} '{cell_text}': expected green, got {run_color}")
                        elif is_negative:
                            # Expect red #C62828
                            if run_color is not None and color_distance(run_color, 0xC6, 0x28, 0x28) < 20:
                                color_ok_count += 1
                            else:
                                color_errors.append(f"Row {row_idx+1} '{cell_text}': expected red, got {run_color}")

                if not color_errors and color_ok_count > 0:
                    print(f"PASS: Component 5 — MoM colors: {color_ok_count} cells colored correctly (0.10 pts)")
                    total_score += 0.10
                else:
                    if color_errors:
                        print(f"FAIL: Component 5 — MoM color errors: {color_errors}")
                    else:
                        print(f"FAIL: Component 5 — No colored MoM cells found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -----------------------------------------------------------------------
    # Component 6: Summary metrics are bold at bottom (0.10 pts)
    # "Total Reach" and "Total Impressions" paragraphs must exist and be bold.
    # In initial state: these texts exist but are NOT bold (size 11pt, no bold).
    # -----------------------------------------------------------------------
    try:
        reach_bold = False
        impressions_bold = False
        for para in doc.paragraphs:
            text = para.text.strip()
            all_runs = [r for r in para.runs if r.text.strip()]
            if 'Total Reach' in text and all_runs:
                if all(r.bold is True for r in all_runs):
                    reach_bold = True
            if 'Total Impressions' in text and all_runs:
                if all(r.bold is True for r in all_runs):
                    impressions_bold = True

        if reach_bold and impressions_bold:
            print(f"PASS: Component 6 — Summary metrics (Total Reach, Total Impressions) are bold (0.10 pts)")
            total_score += 0.10
        else:
            issues = []
            if not reach_bold:
                issues.append("'Total Reach' paragraph not bold")
            if not impressions_bold:
                issues.append("'Total Impressions' paragraph not bold")
            print(f"FAIL: Component 6 — Summary metrics not properly bold: {'; '.join(issues)}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # -----------------------------------------------------------------------
    # Final score
    # -----------------------------------------------------------------------
    final_score = min(round(total_score, 4), 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: run against the canonical file path
file_path = os.path.join(WORKDIR, FILE_NAME)
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
