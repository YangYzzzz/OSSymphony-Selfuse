"""
Initial Setup: Convert footnote 2 on page 3 to endnote
Task ID: writer_struct_031
Domain: libreoffice_writer

Creates a 5-page essay on the Industrial Revolution with 3 footnotes.
Footnote 2 (on page 3) contains the Thompson reference.
No endnotes in initial state.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_031'
OUTPUT = f'{WORKDIR}/history_essay.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_footnote(paragraph, footnote_id, footnote_text):
    """
    Add a footnote reference mark in the paragraph, and add the footnote
    content into the document's footnotes part via XML manipulation.
    """
    doc = paragraph._element.getroottree().getroot()

    # Add footnote reference run in the paragraph
    run = paragraph.add_run()
    run_elem = run._element

    # Create rPr with rStyle for footnote reference
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    run_elem.insert(0, rPr)

    # Create footnote reference element
    footnoteRef = OxmlElement('w:footnoteReference')
    footnoteRef.set(qn('w:id'), str(footnote_id))
    run_elem.append(footnoteRef)

    return run_elem


def ensure_footnotes_part(doc):
    """
    Ensure the document has a footnotes part. Returns the footnotes XML element.
    If it doesn't exist, create it.
    """
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import zipfile
    import io

    # Check if footnotes part already exists
    try:
        footnotes_part = doc.part.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        )
        return footnotes_part
    except Exception:
        pass

    # We need to create the footnotes part
    # Build minimal footnotes XML with separator footnotes
    footnotes_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
             xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"
             xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
             xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink"
             xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d"
             xmlns:o="urn:schemas-microsoft-com:office:office"
             xmlns:oel="http://schemas.microsoft.com/office/2019/extlst"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
             xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
             xmlns:v="urn:schemas-microsoft-com:vml"
             xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
             xmlns:w10="urn:schemas-microsoft-com:office:word"
             xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
             xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
             xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"
             xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"
             xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml"
             xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash"
             xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"
             xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
             xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
             xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
             xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
             mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh wp14">
  <w:footnote w:type="separator" w:id="-1">
    <w:p>
      <w:r>
        <w:separator/>
      </w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p>
      <w:r>
        <w:continuationSeparator/>
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>'''

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
    partname = PackURI('/word/footnotes.xml')
    footnotes_part = Part(partname, content_type, footnotes_xml.encode('utf-8'), doc.part.package)

    rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
    doc.part.relate_to(footnotes_part, rel_type)

    return footnotes_part


def add_footnote_to_part(footnotes_part, footnote_id, footnote_text):
    """Add a footnote entry to the footnotes part XML."""
    from lxml import etree

    # Parse existing XML
    xml_bytes = footnotes_part.blob
    root = etree.fromstring(xml_bytes)

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Build new footnote element
    footnote_elem = etree.SubElement(root, f'{{{W_NS}}}footnote')
    footnote_elem.set(f'{{{W_NS}}}id', str(footnote_id))

    para = etree.SubElement(footnote_elem, f'{{{W_NS}}}p')

    # Paragraph properties with footnote text style
    pPr = etree.SubElement(para, f'{{{W_NS}}}pPr')
    pStyle = etree.SubElement(pPr, f'{{{W_NS}}}pStyle')
    pStyle.set(f'{{{W_NS}}}val', 'FootnoteText')

    # Run with footnote mark
    run_mark = etree.SubElement(para, f'{{{W_NS}}}r')
    rPr_mark = etree.SubElement(run_mark, f'{{{W_NS}}}rPr')
    rStyle_mark = etree.SubElement(rPr_mark, f'{{{W_NS}}}rStyle')
    rStyle_mark.set(f'{{{W_NS}}}val', 'FootnoteReference')
    etree.SubElement(run_mark, f'{{{W_NS}}}footnoteRef')

    # Run with the footnote text
    run_text = etree.SubElement(para, f'{{{W_NS}}}r')
    t_elem = etree.SubElement(run_text, f'{{{W_NS}}}t')
    t_elem.text = footnote_text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Update blob
    footnotes_part._blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('The Industrial Revolution: Transformation of Society and Economy', level=1)

    # Page 1 - Introduction
    intro_heading = doc.add_heading('1. Introduction', level=2)

    p1 = doc.add_paragraph(
        'The Industrial Revolution, which swept through Britain in the late eighteenth and early nineteenth centuries, '
        'represents one of the most profound transformations in human history. Beginning approximately in the 1760s and '
        'continuing through the 1840s, this period witnessed the transition from agrarian, handicraft-based economies to '
        'manufacturing industries and mass production. The consequences of this revolution were felt not only in economic '
        'terms but also in the social, political, and cultural fabric of society.'
    )

    p2 = doc.add_paragraph(
        'Historians have long debated the causes, chronology, and consequences of the Industrial Revolution. Some scholars '
        'emphasize the role of technological innovation—the steam engine, spinning jenny, and power loom—as the primary '
        'drivers of change. Others point to institutional factors such as property rights, financial markets, and colonial '
        'trade networks. Still others highlight the importance of energy resources, particularly coal, in enabling sustained '
        'industrial growth.'
    )

    # Add footnote 1 reference in this paragraph
    fn1_para = doc.add_paragraph(
        'The textile industry was among the first sectors to undergo mechanization. By the 1780s, the spinning jenny and '
        'water frame had revolutionized cotton production, displacing thousands of cottage industry workers and concentrating '
        'production in purpose-built mills along river valleys in Lancashire and Yorkshire.'
    )
    fn1_run = fn1_para.add_run()
    fn1_run_elem = fn1_run._element
    rPr1 = OxmlElement('w:rPr')
    rStyle1 = OxmlElement('w:rStyle')
    rStyle1.set(qn('w:val'), 'FootnoteReference')
    rPr1.append(rStyle1)
    fn1_run_elem.insert(0, rPr1)
    fnRef1 = OxmlElement('w:footnoteReference')
    fnRef1.set(qn('w:id'), '1')
    fn1_run_elem.append(fnRef1)

    doc.add_page_break()

    # Page 2 - Economic Changes
    doc.add_heading('2. Economic Transformation', level=2)

    p3 = doc.add_paragraph(
        'The economic dimensions of the Industrial Revolution were staggering. Britain\'s GDP per capita, which had remained '
        'relatively stagnant for centuries, began to grow at unprecedented rates. Between 1760 and 1840, real wages for '
        'industrial workers rose by approximately 30 percent, though this figure masks significant regional and occupational '
        'variation. Factory owners and merchants accumulated vast fortunes, giving rise to a new industrial bourgeoisie that '
        'would eventually challenge the political dominance of the landed aristocracy.'
    )

    p4 = doc.add_paragraph(
        'The development of canals, turnpike roads, and later railways transformed the geography of commerce. Raw materials '
        'could be transported cheaply from distant regions to industrial centres, while finished goods flowed back to markets '
        'across the country and around the world. By 1850, Britain was responsible for approximately half of the world\'s '
        'iron and cotton cloth production, earning its reputation as the "workshop of the world."'
    )

    p5 = doc.add_paragraph(
        'Banking and credit institutions evolved rapidly to meet the demands of industrial expansion. Joint-stock companies, '
        'insurance markets, and sophisticated bill-discounting networks provided the capital necessary for large-scale '
        'investment in machinery and infrastructure. The Bank of England, established in 1694, played a crucial role in '
        'stabilizing currency and providing credit to both government and commercial enterprises throughout this period of '
        'rapid growth.'
    )

    doc.add_page_break()

    # Page 3 - Social Consequences
    doc.add_heading('3. Social Consequences', level=2)

    p6 = doc.add_paragraph(
        'The social consequences of industrialization were deeply ambivalent. While the long-term trajectory pointed toward '
        'higher living standards for the majority of the population, the short-term experience of industrial workers was '
        'often characterized by poverty, dislocation, and hardship. The factory system imposed new disciplines of time and '
        'work that were alien to workers accustomed to the rhythms of agricultural or artisanal labor.'
    )

    p7 = doc.add_paragraph(
        'The working-class communities that emerged around industrial centres developed their own distinctive cultures, '
        'political traditions, and forms of collective organization. Trade unions, friendly societies, and Chartist movements '
        'represented attempts by workers to assert their interests against the power of employers and the state. These '
        'organizations played a vital role in shaping the political landscape of Victorian Britain and laying the foundations '
        'for the modern labor movement.'
    )

    # This paragraph contains the footnote 2 reference (Thompson reference)
    p8 = doc.add_paragraph(
        'The experience of industrialization was not uniform across the working class. Skilled artisans in trades such as '
        'printing, furniture making, and engineering often maintained considerable autonomy and relatively high wages, even '
        'as unskilled laborers in factories and mines faced deteriorating conditions. The formation of working-class identity '
        'was thus a complex process shaped by occupation, region, gender, and religious affiliation.'
    )
    fn2_run = p8.add_run()
    fn2_run_elem = fn2_run._element
    rPr2 = OxmlElement('w:rPr')
    rStyle2 = OxmlElement('w:rStyle')
    rStyle2.set(qn('w:val'), 'FootnoteReference')
    rPr2.append(rStyle2)
    fn2_run_elem.insert(0, rPr2)
    fnRef2 = OxmlElement('w:footnoteReference')
    fnRef2.set(qn('w:id'), '2')
    fn2_run_elem.append(fnRef2)

    p9 = doc.add_paragraph(
        'Child labor was among the most controversial aspects of early industrialization. Children as young as five or six '
        'worked in textile mills, coal mines, and chimney sweeping trades under conditions that contemporaries increasingly '
        'recognized as morally unacceptable. The Factory Acts of 1833, 1844, and 1847 gradually restricted child labor and '
        'established the principle of state regulation of working conditions, representing a significant departure from '
        'laissez-faire economic philosophy.'
    )

    doc.add_page_break()

    # Page 4 - Technological Innovation
    doc.add_heading('4. Technological Innovation', level=2)

    p10 = doc.add_paragraph(
        'At the heart of the Industrial Revolution was a series of interconnected technological innovations that transformed '
        'productive capacity across multiple industries. The development of the steam engine by James Watt in the 1760s and '
        '1770s was perhaps the most consequential of these innovations, providing a versatile and portable source of '
        'mechanical power that could be applied to virtually any industrial process.'
    )

    p11 = doc.add_paragraph(
        'The iron and steel industries underwent parallel revolutions in technology. Abraham Darby\'s development of coke '
        'smelting at Coalbrookdale in 1709 had laid the groundwork for a dramatic expansion of iron production. By the early '
        'nineteenth century, puddling and rolling techniques allowed for the mass production of wrought iron, which found '
        'applications in everything from machinery components to bridge construction and eventually the girders and rails of '
        'the expanding railway network.'
    )

    p12 = doc.add_paragraph(
        'The chemical industry, though less visible than textiles or iron, also underwent significant transformation during '
        'this period. The Leblanc process for manufacturing soda ash, developed in France in the 1790s, was quickly adopted '
        'in Britain and provided essential inputs for the glass, soap, and textile bleaching industries. Chlorine bleaching, '
        'developed by Claude Berthollet in 1785, dramatically reduced the time and land required to bleach cotton cloth, '
        'enabling textile manufacturers to keep pace with the growing output of mechanized spinning mills.'
    )

    doc.add_page_break()

    # Page 5 - Conclusion
    doc.add_heading('5. Legacy and Conclusion', level=2)

    p13 = doc.add_paragraph(
        'The Industrial Revolution fundamentally reshaped the world in ways that continue to echo through the present day. '
        'The patterns of urbanization, class formation, environmental change, and global trade established during this period '
        'created the structural foundations of modern industrial society. The environmental consequences of coal-based '
        'industrialization, from air and water pollution to carbon emissions that are now recognized as drivers of climate '
        'change, represent a profound and continuing legacy of this transformative period.'
    )

    p14 = doc.add_paragraph(
        'The spread of industrialization from Britain to continental Europe and North America during the nineteenth century, '
        'and subsequently to Asia, Latin America, and Africa, created a globalized industrial economy with massive inequalities '
        'between and within nations. The mechanisms by which some nations industrialized successfully while others remained '
        'primary commodity exporters have been subjects of intense scholarly debate, with implications for contemporary '
        'development policy and practice.'
    )

    p15 = doc.add_paragraph(
        'In conclusion, the Industrial Revolution represents a watershed in human history, marking the transition from '
        'pre-modern to modern patterns of production, consumption, and social organization. Understanding its causes, '
        'dynamics, and consequences remains essential for historians, economists, and policymakers seeking to address the '
        'challenges of contemporary industrial society. The debates initiated by early social historians about the costs '
        'and benefits of industrialization for working people continue to resonate in present-day discussions of economic '
        'development, inequality, and technological change.'
    )

    # Add footnote 3 reference
    p16 = doc.add_paragraph(
        'Future research directions include comparative studies of industrialization across different national contexts, '
        'greater attention to the experiences of women and marginalized groups within industrial society, and interdisciplinary '
        'approaches that integrate economic, social, cultural, and environmental history. The Industrial Revolution, far from '
        'being a closed chapter, remains an open and vital field of historical inquiry.'
    )
    fn3_run = p16.add_run()
    fn3_run_elem = fn3_run._element
    rPr3 = OxmlElement('w:rPr')
    rStyle3 = OxmlElement('w:rStyle')
    rStyle3.set(qn('w:val'), 'FootnoteReference')
    rPr3.append(rStyle3)
    fn3_run_elem.insert(0, rPr3)
    fnRef3 = OxmlElement('w:footnoteReference')
    fnRef3.set(qn('w:id'), '3')
    fn3_run_elem.append(fnRef3)

    # Save the document first to have a valid zip/package
    doc.save(OUTPUT)

    # Now reopen and add footnotes part via direct XML manipulation
    from docx import Document as DocxDocument
    import zipfile
    from lxml import etree
    import shutil

    # Read the saved docx as zip and add footnotes.xml
    temp_output = OUTPUT + '.tmp'
    shutil.copy(OUTPUT, temp_output)

    footnotes_xml_content = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
             xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p>
      <w:r>
        <w:separator/>
      </w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p>
      <w:r>
        <w:continuationSeparator/>
      </w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:id="1">
    <w:p>
      <w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>
      <w:r>
        <w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r><w:t xml:space="preserve">Mokyr, Joel, The Enlightened Economy: An Economic History of Britain 1700-1850, Yale University Press, 2009, p. 47.</w:t></w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:id="2">
    <w:p>
      <w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>
      <w:r>
        <w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r><w:t xml:space="preserve">Thompson, E.P., The Making of the English Working Class, 1963, p. 214.</w:t></w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:id="3">
    <w:p>
      <w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>
      <w:r>
        <w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>
        <w:footnoteRef/>
      </w:r>
      <w:r><w:t xml:space="preserve">Berg, Maxine, The Age of Manufactures, 1700-1820: Industry, Innovation and Work in Britain, Routledge, 1994, p. 312.</w:t></w:r>
    </w:p>
  </w:footnote>
</w:footnotes>'''

    with zipfile.ZipFile(temp_output, 'r') as zin:
        with zipfile.ZipFile(OUTPUT, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/_rels/document.xml.rels':
                    # Add footnotes relationship
                    tree = etree.fromstring(data)
                    ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
                    # Find next rel id
                    existing_ids = [el.get('Id') for el in tree.findall(f'{{{ns}}}Relationship')]
                    max_id = 0
                    for eid in existing_ids:
                        if eid and eid.startswith('rId'):
                            try:
                                max_id = max(max_id, int(eid[3:]))
                            except ValueError:
                                pass
                    new_rel_id = f'rId{max_id + 1}'
                    new_rel = etree.SubElement(tree, f'{{{ns}}}Relationship')
                    new_rel.set('Id', new_rel_id)
                    new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
                    new_rel.set('Target', 'footnotes.xml')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
                zout.writestr(item, data)
            # Add the footnotes.xml file
            zout.writestr('word/footnotes.xml', footnotes_xml_content.encode('utf-8'))

    os.remove(temp_output)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
