"""
Reward Script: Set up different headers for odd and even pages
Task ID: writer_hr_047
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25) - evenAndOddHeaders enabled in document settings
  Component 2 (0.30) - Odd (default) header text is 'Employee Handbook 2026'
  Component 3 (0.15) - Odd header is RIGHT aligned
  Component 4 (0.20) - Even header text is 'Meridian Corp - Human Resources'
  Component 5 (0.10) - Even header is LEFT aligned
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_047'


def persist_app_state(domain: str):
    """Try to save any unsaved changes in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: evenAndOddHeaders enabled in document settings (0.25 points)
    # This is the master toggle that enables different headers on odd/even pages.
    try:
        settings_elem = doc.settings.element
        even_odd_elem = settings_elem.find(qn('w:evenAndOddHeaders'))
        if even_odd_elem is not None:
            # The element exists. In OOXML, presence means enabled unless val="false"/"0".
            val = even_odd_elem.get(qn('w:val'))
            if val is None or val.lower() in ('true', '1', 'on'):
                print(f"PASS: Component 1 -- evenAndOddHeaders is enabled (val={val}) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 -- evenAndOddHeaders exists but val={val}")
        else:
            print("FAIL: Component 1 -- evenAndOddHeaders element not found in document settings")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Odd (default) header text is 'Employee Handbook 2026' (0.30 points)
    # In OOXML, the 'default' header is the odd-page header.
    try:
        section = doc.sections[0]
        header = section.header
        header_text = ""
        if header.paragraphs:
            header_text = header.paragraphs[0].text.strip()
        if header_text == "Employee Handbook 2026":
            print(f"PASS: Component 2 -- Odd header text = '{header_text}' (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 -- Expected 'Employee Handbook 2026', found '{header_text}'")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Odd header is RIGHT aligned (0.15 points)
    try:
        section = doc.sections[0]
        header = section.header
        if header.paragraphs:
            alignment = header.paragraphs[0].paragraph_format.alignment
            if alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                print(f"PASS: Component 3 -- Odd header is RIGHT aligned (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 -- Odd header alignment = {alignment}, expected RIGHT (2)")
        else:
            print("FAIL: Component 3 -- No paragraphs in odd header")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Even header text is 'Meridian Corp - Human Resources' (0.20 points)
    try:
        section = doc.sections[0]
        even_header = section.even_page_header
        even_text = ""
        if even_header.paragraphs:
            even_text = even_header.paragraphs[0].text.strip()
        if even_text == "Meridian Corp - Human Resources":
            print(f"PASS: Component 4 -- Even header text = '{even_text}' (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 -- Expected 'Meridian Corp - Human Resources', found '{even_text}'")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: Even header is LEFT aligned (0.10 points)
    try:
        section = doc.sections[0]
        even_header = section.even_page_header
        if even_header.paragraphs:
            alignment = even_header.paragraphs[0].paragraph_format.alignment
            # LEFT alignment can be explicit (0) or None (default is LEFT)
            if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT or alignment is None:
                # Need to distinguish: in golden, it should be explicitly LEFT.
                # But we also need initial to fail. In initial, even_header is empty text.
                # So we gate this on the even header having the correct text too.
                even_text = even_header.paragraphs[0].text.strip()
                if even_text == "Meridian Corp - Human Resources":
                    print(f"PASS: Component 5 -- Even header is LEFT aligned (alignment={alignment}) (0.10 pts)")
                    total_score += 0.10
                else:
                    print(f"FAIL: Component 5 -- Even header text mismatch, cannot award alignment points")
            else:
                print(f"FAIL: Component 5 -- Even header alignment = {alignment}, expected LEFT (0)")
        else:
            print("FAIL: Component 5 -- No paragraphs in even header")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
