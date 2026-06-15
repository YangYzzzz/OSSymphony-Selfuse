"""
Initial Setup: Create a 20-page employee handbook with no headers configured.
Task ID: writer_hr_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_047'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # NO headers configured -- this is intentional for the task

    # === Title Page ===
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading('Employee Handbook 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Corp')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run('Human Resources Department')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_p.add_run('Effective January 1, 2026')
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_page_break()

    # === Table of Contents ===
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Welcome to Meridian Corp',
        '2. Company Overview and Mission',
        '3. Employment Policies',
        '4. Compensation and Benefits',
        '5. Workplace Conduct',
        '6. Leave Policies',
        '7. Performance Management',
        '8. Health and Safety',
        '9. Information Security',
        '10. Employee Development',
        '11. Separation of Employment',
        '12. Acknowledgment',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # === Chapter 1: Welcome ===
    add_heading(doc, '1. Welcome to Meridian Corp', level=1)
    add_body(doc, (
        'Welcome to Meridian Corp! We are excited to have you as part of our team. '
        'This handbook has been prepared to provide you with a comprehensive overview of '
        'the company\'s policies, procedures, and benefits. We encourage you to read it carefully '
        'and keep it accessible for future reference.'
    ))
    add_body(doc, (
        'At Meridian Corp, we believe that our people are our greatest asset. Since our '
        'founding in 1998 by CEO Alexandra Whitfield, we have grown from a small consulting '
        'firm into a global technology leader with over 4,500 employees across 12 offices '
        'in North America, Europe, and Asia-Pacific.'
    ))
    add_body(doc, (
        'Our success is built on a culture of innovation, respect, and collaboration. '
        'Every team member, regardless of role or seniority, plays a vital part in achieving '
        'our vision: to transform how businesses leverage technology to create value and drive '
        'sustainable growth.'
    ))
    add_body(doc, (
        'This handbook is not a contract of employment, nor is it intended to create any '
        'contractual obligations. Meridian Corp reserves the right to modify policies and '
        'procedures described herein at any time, with appropriate notice to employees.'
    ))

    doc.add_page_break()

    # === Chapter 2: Company Overview ===
    add_heading(doc, '2. Company Overview and Mission', level=1)
    add_heading(doc, '2.1 Our Mission', level=2)
    add_body(doc, (
        'Meridian Corp\'s mission is to empower organizations through innovative technology '
        'solutions that drive measurable business outcomes. We are committed to delivering '
        'excellence in every engagement, building lasting partnerships with our clients, '
        'and fostering a workplace where talented individuals can thrive and grow.'
    ))
    add_heading(doc, '2.2 Core Values', level=2)
    add_bullet_list(doc, [
        'Innovation: We continuously push boundaries to find creative solutions.',
        'Integrity: We conduct business with transparency and honesty.',
        'Collaboration: We achieve more together than we ever could alone.',
        'Excellence: We hold ourselves to the highest standards of quality.',
        'Respect: We value diverse perspectives and treat everyone with dignity.',
    ])
    add_heading(doc, '2.3 Organizational Structure', level=2)
    add_body(doc, (
        'Meridian Corp operates through five core divisions: Technology Solutions, '
        'Enterprise Consulting, Cloud Services, Data Analytics, and Digital Transformation. '
        'Each division is led by a Senior Vice President who reports directly to the '
        'Chief Executive Officer. Regional operations are managed through our global '
        'office network, with headquarters located in San Francisco, California.'
    ))
    add_body(doc, (
        'The Human Resources department supports all divisions and is responsible for '
        'talent acquisition, employee relations, benefits administration, learning and '
        'development, and organizational effectiveness. HR operates under the leadership '
        'of Chief People Officer Dr. Rachel Okonkwo.'
    ))

    doc.add_page_break()

    # === Chapter 3: Employment Policies ===
    add_heading(doc, '3. Employment Policies', level=1)
    add_heading(doc, '3.1 Equal Employment Opportunity', level=2)
    add_body(doc, (
        'Meridian Corp is an equal opportunity employer. We do not discriminate on the basis '
        'of race, color, religion, sex, national origin, age, disability, genetic information, '
        'veteran status, sexual orientation, gender identity, or any other characteristic '
        'protected by applicable law. This policy applies to all terms and conditions of '
        'employment, including recruitment, hiring, placement, promotion, termination, layoff, '
        'recall, transfer, leaves of absence, compensation, and training.'
    ))
    add_heading(doc, '3.2 Employment Classification', level=2)
    add_body(doc, (
        'Employees are classified as either full-time (40 hours per week), part-time '
        '(20-39 hours per week), or temporary. Full-time and eligible part-time employees '
        'receive the complete benefits package described in Chapter 4. Temporary employees '
        'may be eligible for limited benefits as outlined in their offer letters.'
    ))
    add_heading(doc, '3.3 Probationary Period', level=2)
    add_body(doc, (
        'All new employees undergo a 90-day probationary period during which performance, '
        'attendance, and cultural fit are evaluated. During this period, either the employee '
        'or Meridian Corp may end the employment relationship with two weeks\' notice. '
        'Upon successful completion of the probationary period, employees transition to '
        'regular employment status.'
    ))
    add_heading(doc, '3.4 Background Checks', level=2)
    add_body(doc, (
        'All offers of employment at Meridian Corp are contingent upon the successful '
        'completion of a background check, which may include verification of employment '
        'history, education credentials, criminal record check, and professional reference checks. '
        'The scope of the background check may vary depending on the position and applicable laws.'
    ))

    doc.add_page_break()

    # === Chapter 4: Compensation and Benefits ===
    add_heading(doc, '4. Compensation and Benefits', level=1)
    add_heading(doc, '4.1 Compensation Philosophy', level=2)
    add_body(doc, (
        'Meridian Corp is committed to providing competitive compensation that attracts '
        'and retains top talent. Our compensation structure is reviewed annually and '
        'benchmarked against industry standards. Base salaries are determined by role, '
        'experience, performance, and market data.'
    ))
    add_heading(doc, '4.2 Pay Schedule', level=2)
    add_body(doc, (
        'Employees are paid on a bi-weekly basis, with 26 pay periods per calendar year. '
        'Payday is every other Friday. Direct deposit is available and strongly encouraged. '
        'Pay stubs can be accessed through the Meridian Corp employee portal at '
        'portal.meridiancorp.com.'
    ))
    add_heading(doc, '4.3 Benefits Overview', level=2)
    add_body(doc, 'Meridian Corp offers a comprehensive benefits package, including:')
    add_bullet_list(doc, [
        'Medical, dental, and vision insurance (Blue Cross Blue Shield PPO and HMO options)',
        '401(k) retirement plan with 6% company match (vested after 2 years)',
        'Life insurance and accidental death & dismemberment (AD&D) coverage',
        'Short-term and long-term disability insurance',
        'Flexible Spending Account (FSA) and Health Savings Account (HSA)',
        'Employee Assistance Program (EAP) providing confidential counseling',
        'Tuition reimbursement up to $5,250 per calendar year',
        'Commuter benefits and parking subsidies',
        'Gym membership discount (up to 50% at partner facilities)',
    ])
    add_heading(doc, '4.4 Overtime', level=2)
    add_body(doc, (
        'Non-exempt employees are eligible for overtime pay at 1.5 times their regular '
        'hourly rate for hours worked in excess of 40 per workweek. All overtime must be '
        'pre-approved by the employee\'s direct supervisor. Unauthorized overtime may result '
        'in disciplinary action.'
    ))

    doc.add_page_break()

    # === Chapter 5: Workplace Conduct ===
    add_heading(doc, '5. Workplace Conduct', level=1)
    add_heading(doc, '5.1 Code of Conduct', level=2)
    add_body(doc, (
        'Employees are expected to conduct themselves in a professional manner at all times. '
        'This includes treating colleagues, clients, vendors, and visitors with respect and '
        'courtesy. Meridian Corp has zero tolerance for harassment, bullying, discrimination, '
        'or retaliation of any kind.'
    ))
    add_heading(doc, '5.2 Anti-Harassment Policy', level=2)
    add_body(doc, (
        'Meridian Corp prohibits all forms of harassment, including but not limited to '
        'sexual harassment, racial harassment, and harassment based on any protected characteristic. '
        'Harassment includes unwelcome verbal, physical, or visual conduct that creates an '
        'intimidating, hostile, or offensive work environment.'
    ))
    add_body(doc, (
        'Employees who experience or witness harassment should report it immediately to their '
        'supervisor, HR Business Partner, or through the anonymous ethics hotline at '
        '1-800-555-0147. All complaints are investigated promptly and confidentially. '
        'Retaliation against anyone who reports harassment in good faith is strictly prohibited '
        'and will result in disciplinary action up to and including termination.'
    ))
    add_heading(doc, '5.3 Dress Code', level=2)
    add_body(doc, (
        'Meridian Corp maintains a business casual dress code Monday through Thursday. '
        'Casual attire is permitted on Fridays. Employees in client-facing roles may be '
        'required to adhere to a more formal dress code when meeting with clients. '
        'All employees should use good judgment and dress in a manner appropriate '
        'for their work environment and scheduled activities.'
    ))
    add_heading(doc, '5.4 Attendance and Punctuality', level=2)
    add_body(doc, (
        'Regular attendance and punctuality are essential. Employees are expected to arrive '
        'on time and adhere to their scheduled work hours. If an employee will be late or '
        'absent, they must notify their supervisor as soon as possible, ideally before the '
        'start of their shift. Excessive absenteeism or tardiness may result in disciplinary '
        'action as outlined in Section 5.6.'
    ))
    add_heading(doc, '5.5 Substance Abuse', level=2)
    add_body(doc, (
        'Meridian Corp maintains a drug-free workplace. The use, possession, sale, or '
        'distribution of illegal drugs or controlled substances on company premises or '
        'during work hours is strictly prohibited. Employees who are under the influence '
        'of drugs or alcohol while at work will be subject to immediate disciplinary action. '
        'Employees who voluntarily seek assistance for substance abuse issues will be '
        'supported through the Employee Assistance Program.'
    ))
    add_heading(doc, '5.6 Progressive Discipline', level=2)
    add_body(doc, (
        'Meridian Corp follows a progressive discipline process: (1) verbal warning, '
        '(2) written warning, (3) final written warning with performance improvement plan, '
        'and (4) termination. Depending on the severity of the infraction, steps may be '
        'skipped. Serious offenses such as theft, violence, or gross misconduct may result '
        'in immediate termination without prior warnings.'
    ))

    doc.add_page_break()

    # === Chapter 6: Leave Policies ===
    add_heading(doc, '6. Leave Policies', level=1)
    add_heading(doc, '6.1 Paid Time Off (PTO)', level=2)
    add_body(doc, (
        'Full-time employees accrue PTO based on years of service: 15 days for 0-3 years, '
        '20 days for 4-7 years, and 25 days for 8+ years. PTO can be used for vacation, '
        'personal days, or sick leave. Unused PTO may be carried over up to a maximum of '
        '10 days into the following calendar year. PTO requests should be submitted at least '
        'two weeks in advance for planned absences.'
    ))
    add_heading(doc, '6.2 Holidays', level=2)
    add_body(doc, 'Meridian Corp observes the following paid holidays:')
    add_bullet_list(doc, [
        "New Year's Day (January 1)",
        'Martin Luther King Jr. Day (Third Monday in January)',
        "Presidents' Day (Third Monday in February)",
        'Memorial Day (Last Monday in May)',
        'Independence Day (July 4)',
        'Labor Day (First Monday in September)',
        'Thanksgiving Day (Fourth Thursday in November)',
        'Day After Thanksgiving (Fourth Friday in November)',
        'Christmas Eve (December 24)',
        'Christmas Day (December 25)',
    ])
    add_heading(doc, '6.3 Family and Medical Leave', level=2)
    add_body(doc, (
        'Eligible employees may take up to 12 weeks of unpaid, job-protected leave per year '
        'under the Family and Medical Leave Act (FMLA) for qualifying reasons, including the '
        'birth or adoption of a child, the serious health condition of the employee or an '
        'immediate family member, or qualifying exigency related to a family member\'s military '
        'service. Meridian Corp additionally provides 8 weeks of paid parental leave for the '
        'birth or adoption of a child.'
    ))
    add_heading(doc, '6.4 Bereavement Leave', level=2)
    add_body(doc, (
        'Employees may take up to 5 days of paid bereavement leave for the death of an '
        'immediate family member (spouse, child, parent, sibling, grandparent, or grandchild). '
        'Up to 3 days of paid leave is provided for the death of an extended family member '
        '(aunt, uncle, cousin, in-law). Additional unpaid leave may be granted at the '
        'discretion of the employee\'s manager.'
    ))

    doc.add_page_break()

    # === Chapter 7: Performance Management ===
    add_heading(doc, '7. Performance Management', level=1)
    add_heading(doc, '7.1 Performance Review Cycle', level=2)
    add_body(doc, (
        'Meridian Corp conducts annual performance reviews in Q1 of each calendar year, '
        'covering the prior year\'s performance. The review process includes a self-assessment, '
        'manager evaluation, and calibration sessions to ensure consistency and fairness '
        'across the organization. Mid-year check-ins are conducted in July to track progress '
        'on goals and provide ongoing feedback.'
    ))
    add_heading(doc, '7.2 Goal Setting', level=2)
    add_body(doc, (
        'Each employee works with their manager to establish SMART goals (Specific, Measurable, '
        'Achievable, Relevant, Time-bound) at the beginning of each year. Goals should align '
        'with departmental objectives and the company\'s strategic priorities. Progress toward '
        'goals is reviewed during mid-year check-ins and the annual performance review.'
    ))
    add_heading(doc, '7.3 Performance Ratings', level=2)
    add_body(doc, (
        'Employees are rated on a five-point scale: (5) Exceptional - consistently exceeds '
        'expectations in all areas; (4) Exceeds Expectations - regularly exceeds expectations '
        'in most areas; (3) Meets Expectations - consistently meets job requirements; '
        '(2) Needs Improvement - occasionally falls below expectations; (1) Unsatisfactory - '
        'fails to meet minimum job requirements. Performance ratings directly impact annual '
        'merit increases and bonus eligibility.'
    ))
    add_heading(doc, '7.4 Career Development', level=2)
    add_body(doc, (
        'Meridian Corp encourages employees to take ownership of their career development. '
        'Resources include the internal learning management system (Meridian Academy), '
        'mentorship programs, leadership development tracks, and tuition reimbursement. '
        'Employees are encouraged to discuss career aspirations with their managers during '
        'regular one-on-one meetings and performance reviews.'
    ))

    doc.add_page_break()

    # === Chapter 8: Health and Safety ===
    add_heading(doc, '8. Health and Safety', level=1)
    add_heading(doc, '8.1 Workplace Safety', level=2)
    add_body(doc, (
        'Meridian Corp is committed to providing a safe and healthy work environment for '
        'all employees, visitors, and contractors. The company complies with all applicable '
        'Occupational Safety and Health Administration (OSHA) regulations and maintains '
        'comprehensive safety protocols across all office locations.'
    ))
    add_heading(doc, '8.2 Emergency Procedures', level=2)
    add_body(doc, (
        'Each office location maintains an Emergency Action Plan (EAP) that includes '
        'evacuation routes, assembly points, and emergency contact information. Fire drills '
        'are conducted quarterly. All employees are required to familiarize themselves with '
        'the EAP for their office location and participate in all scheduled drills. '
        'Floor wardens are designated for each floor and trained in emergency response procedures.'
    ))
    add_heading(doc, '8.3 Workplace Ergonomics', level=2)
    add_body(doc, (
        'Meridian Corp provides ergonomic assessments for all workstations upon request. '
        'Standing desks, ergonomic chairs, monitor arms, and other equipment are available '
        'through the Facilities team. Employees are encouraged to take regular breaks, '
        'practice proper posture, and report any discomfort or repetitive strain symptoms '
        'to their manager or the HR team.'
    ))
    add_heading(doc, '8.4 Workers\' Compensation', level=2)
    add_body(doc, (
        'Employees who are injured or become ill as a result of their job duties are covered '
        'under workers\' compensation insurance. All work-related injuries or illnesses must be '
        'reported to a supervisor immediately. The company will work with the employee and '
        'healthcare providers to facilitate a safe and timely return to work.'
    ))

    doc.add_page_break()

    # === Chapter 9: Information Security ===
    add_heading(doc, '9. Information Security', level=1)
    add_heading(doc, '9.1 Data Protection', level=2)
    add_body(doc, (
        'All employees are responsible for protecting Meridian Corp\'s proprietary information '
        'and client data. This includes complying with the company\'s Information Security Policy, '
        'using strong passwords (minimum 12 characters with complexity requirements), enabling '
        'multi-factor authentication on all company accounts, and reporting any suspected '
        'security incidents to the IT Security team at security@meridiancorp.com.'
    ))
    add_heading(doc, '9.2 Acceptable Use Policy', level=2)
    add_body(doc, (
        'Company-provided technology resources, including computers, mobile devices, email, '
        'and internet access, are primarily for business use. Limited personal use is permitted '
        'provided it does not interfere with job duties, violate company policies, or pose '
        'security risks. Employees should have no expectation of privacy when using company '
        'technology resources.'
    ))
    add_heading(doc, '9.3 Remote Work Security', level=2)
    add_body(doc, (
        'Employees working remotely must use the company VPN when accessing internal systems. '
        'Work should be performed on company-issued devices only. Employees must ensure their '
        'home network is secured with WPA3 encryption and that sensitive documents are stored '
        'in approved cloud services (Microsoft 365 or Google Workspace) rather than on local drives.'
    ))
    add_heading(doc, '9.4 Confidentiality Agreements', level=2)
    add_body(doc, (
        'All employees sign a confidentiality and non-disclosure agreement (NDA) upon hire. '
        'This agreement remains in effect during employment and for a period of two years '
        'following separation. Employees must not disclose proprietary information, trade '
        'secrets, client data, or internal business strategies to unauthorized individuals '
        'or organizations.'
    ))

    doc.add_page_break()

    # === Chapter 10: Employee Development ===
    add_heading(doc, '10. Employee Development', level=1)
    add_heading(doc, '10.1 Learning and Development Programs', level=2)
    add_body(doc, (
        'Meridian Corp invests significantly in employee development through Meridian Academy, '
        'our internal learning platform. The Academy offers over 500 courses spanning technical '
        'skills, leadership development, project management, and soft skills. All employees '
        'are encouraged to complete at least 40 hours of professional development annually.'
    ))
    add_heading(doc, '10.2 Mentorship Program', level=2)
    add_body(doc, (
        'The Meridian Mentorship Program pairs junior employees with experienced leaders '
        'for a 12-month engagement. Mentors and mentees meet at least twice per month to '
        'discuss career goals, professional challenges, and skill development. The program '
        'has a 92% satisfaction rate and has been instrumental in promoting internal talent mobility.'
    ))
    add_heading(doc, '10.3 Tuition Reimbursement', level=2)
    add_body(doc, (
        'Full-time employees who have completed their probationary period are eligible for '
        'tuition reimbursement of up to $5,250 per calendar year for courses related to their '
        'current role or career path at Meridian Corp. Courses must be taken at accredited '
        'institutions and a minimum grade of B is required for reimbursement. Approval from '
        'the employee\'s manager and HR is required prior to enrollment.'
    ))
    add_heading(doc, '10.4 Internal Mobility', level=2)
    add_body(doc, (
        'Meridian Corp encourages internal mobility and promotes from within whenever possible. '
        'All open positions are posted on the internal job board for a minimum of 5 business '
        'days before external advertising. Employees in good standing who have been in their '
        'current role for at least 12 months are eligible to apply for internal transfers. '
        'The hiring manager and HR will evaluate internal candidates based on qualifications, '
        'performance history, and career development goals.'
    ))

    doc.add_page_break()

    # === Chapter 11: Separation of Employment ===
    add_heading(doc, '11. Separation of Employment', level=1)
    add_heading(doc, '11.1 Voluntary Resignation', level=2)
    add_body(doc, (
        'Employees who wish to resign from their position at Meridian Corp are requested to '
        'provide a minimum of two weeks\' written notice to their manager. Managers and directors '
        'are requested to provide four weeks\' notice. The resignation letter should include '
        'the employee\'s last intended work date. HR will schedule an exit interview to discuss '
        'the employee\'s experience and collect company property.'
    ))
    add_heading(doc, '11.2 Involuntary Termination', level=2)
    add_body(doc, (
        'Meridian Corp reserves the right to terminate employment at any time, with or without '
        'cause, subject to applicable law. Involuntary terminations may result from performance '
        'issues, policy violations, reorganization, or reduction in force. The HR team ensures '
        'all terminations are handled fairly, consistently, and in compliance with applicable '
        'employment laws.'
    ))
    add_heading(doc, '11.3 Exit Process', level=2)
    add_body(doc, (
        'Upon separation, employees must return all company property, including laptops, '
        'access badges, mobile devices, keys, and documents. Final paychecks are processed '
        'in accordance with state law. Benefits continuation under COBRA will be offered to '
        'eligible employees. Information about 401(k) rollover options and other post-employment '
        'matters will be provided during the exit process.'
    ))

    doc.add_page_break()

    # === Chapter 12: Acknowledgment ===
    add_heading(doc, '12. Acknowledgment', level=1)
    add_body(doc, (
        'By signing below, I acknowledge that I have received and read the Meridian Corp '
        'Employee Handbook 2026. I understand that this handbook provides a general overview '
        'of company policies and is not a contract of employment. I agree to abide by the '
        'policies, procedures, and guidelines described herein.'
    ))
    for _ in range(3):
        doc.add_paragraph()
    add_body(doc, '_' * 50)
    add_body(doc, 'Employee Name (Printed)')
    doc.add_paragraph()
    add_body(doc, '_' * 50)
    add_body(doc, 'Employee Signature')
    doc.add_paragraph()
    add_body(doc, '_' * 50)
    add_body(doc, 'Date')
    doc.add_paragraph()
    doc.add_paragraph()
    add_body(doc, '_' * 50)
    add_body(doc, 'HR Representative Signature')
    doc.add_paragraph()
    add_body(doc, '_' * 50)
    add_body(doc, 'Date')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
