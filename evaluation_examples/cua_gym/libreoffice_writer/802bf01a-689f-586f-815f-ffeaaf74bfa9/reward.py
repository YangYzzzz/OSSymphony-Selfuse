"""
Reward Script: Insert a 4x6 project timeline table with formatted header row
Task ID: writer_rd_012
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with 4 columns and 6 rows
  Component 2 (0.25): Header row contains correct text (Phase, Start Date, End Date, Status)
  Component 3 (0.25): Header row cells have dark blue background (#003366)
  Component 4 (0.25): Header row text is white (#FFFFFF) and bold
"""

import os
from docx import Document
from docx.oxml.ns import qn
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_012'


def color_distance_hex(hex1, hex2):
    """Compute Euclidean distance between two hex color strings (e.g. '003366')."""
    r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
    r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with 4 columns and 6 rows (0.25 points)
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 — No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 6 and num_cols == 4:
                print(f"PASS: Component 1 — Table has {num_rows} rows x {num_cols} cols (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — Expected 6x4 table, found {num_rows}x{num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Gate: need at least one table to check remaining components
    if len(doc.tables) == 0:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]

    # Component 2: Header row text matches expected values (0.25 points)
    try:
        expected_headers = ['Phase', 'Start Date', 'End Date', 'Status']
        actual_headers = []
        if len(table.rows) >= 1:
            for ci, cell in enumerate(table.rows[0].cells):
                actual_headers.append(cell.text.strip())

        matches = 0
        for i, expected in enumerate(expected_headers):
            if i < len(actual_headers) and actual_headers[i].lower() == expected.lower():
                matches += 1

        if matches == 4:
            print(f"PASS: Component 2 — Header text matches: {actual_headers} (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Expected {expected_headers}, found {actual_headers} ({matches}/4 match)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row cells have dark blue background #003366 (0.25 points)
    try:
        TARGET_BG = '003366'
        bg_ok_count = 0

        if len(table.rows) >= 1:
            for ci, cell in enumerate(table.rows[0].cells):
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and color_distance_hex(fill.upper(), TARGET_BG.upper()) < 30:
                            bg_ok_count += 1
                        else:
                            print(f"  Cell(0,{ci}): fill={fill}, expected ~{TARGET_BG}")

        if bg_ok_count == 4:
            print(f"PASS: Component 3 — All 4 header cells have background #003366 (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Only {bg_ok_count}/4 header cells have correct background")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Header row text is white (#FFFFFF) and bold (0.25 points)
    try:
        TARGET_TEXT_COLOR = 'FFFFFF'
        format_ok_count = 0

        if len(table.rows) >= 1:
            for ci, cell in enumerate(table.rows[0].cells):
                # Collect all runs from the cell
                all_runs = [run for para in cell.paragraphs for run in para.runs]
                # Check if any run is bold
                cell_bold = any(run.font.bold for run in all_runs)
                # Check if any run has white text color
                cell_white = any(
                    run.font.color and run.font.color.rgb and
                    color_distance_hex(str(run.font.color.rgb), TARGET_TEXT_COLOR) < 30
                    for run in all_runs
                )

                if cell_bold and cell_white:
                    format_ok_count += 1
                else:
                    print(f"  Cell(0,{ci}): bold={cell_bold}, white={cell_white}")

        if format_ok_count == 4:
            print(f"PASS: Component 4 — All 4 header cells have white bold text (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 — Only {format_ok_count}/4 header cells have correct formatting")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
