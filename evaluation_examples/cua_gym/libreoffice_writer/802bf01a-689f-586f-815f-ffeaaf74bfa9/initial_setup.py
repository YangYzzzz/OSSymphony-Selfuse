"""
Initial Setup: Insert a formatted project timeline table into a Writer document.
Task ID: writer_rd_012
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_012'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font to Liberation Serif 11pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Liberation Serif'
    font.size = Pt(11)

    # Add heading
    heading = doc.add_heading('Project Timeline', level=1)
    for run in heading.runs:
        run.font.name = 'Liberation Serif'

    # Add a blank paragraph (where the table should be inserted by the agent)
    blank_para = doc.add_paragraph('')

    # Add some introductory context below to make the document more realistic
    intro = doc.add_paragraph(
        'The following section is reserved for the detailed project timeline. '
        'Please insert a table below the heading to outline each phase of the project, '
        'including start dates, end dates, and current status.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
