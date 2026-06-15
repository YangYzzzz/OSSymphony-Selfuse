"""
FINAL REWARD SCRIPT - SUCCESS
Task: LibreOffice Writer question: I need the line "Table 2: Metrics" to show up as an official caption right under Table 2 (Category = Table, Position = Below). How do I insert that correctly so it follows the program’s numbering rules?
Generated: 2025-09-10 19:37:08
Status: success
Model: azure-o3
Total Steps: 7
"""

import os
import re
import zipfile
from docx import Document
import lxml.etree as ET


def _extract_document_blocks(file_path):
    """Return a list representing the block-level order of the DOCX body.
    Each element is a tuple: (kind, text) where kind is 'tbl' or 'p'.
    For paragraphs we also return their concatenated text (empty string if none).
    """
    with zipfile.ZipFile(file_path) as z:
        xml_bytes = z.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    blocks = []
    body = root.find(".//w:body", NS)
    for child in body:
        tag = ET.QName(child).localname
        if tag == "tbl":
            blocks.append(("tbl", ""))
        elif tag == "p":
            text = "".join(t.text or "" for t in child.findall(".//w:t", NS))
            blocks.append(("p", text))
    return blocks


def verify_table_caption_task(file_path):
    """Verification for the LibreOffice Writer caption task.

    Requirements interpreted from the prompt:
      1. Document must contain at least two tables.
      2. There must be a caption paragraph with the exact numbering text
         "Table 2: Metrics" (case-insensitive, accepts minor spacing/punctuation variations).
      3. The caption paragraph must use a Caption style (style name containing
         the word 'caption').
      4. The caption must appear immediately below the second table (allowing
         for a single blank paragraph in between).

    Progressive scoring delivers up to 1.0 only when every requirement passes.
    """
    print(f"Starting verification for: {os.path.basename(file_path)}")
    score = 0.0
    max_score = 1.0

    # ---------- Step 0: Load document ----------
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Could not load DOCX: {e}")
        return 0.0  # Fatal – nothing else can be checked

    # ---------- Step 1: Tables existence ----------
    if len(doc.tables) >= 2:
        print(f"✓ Found {len(doc.tables)} tables (need at least 2)")
        score += 0.1  # small partial credit – prerequisite but still user action
    else:
        print("✗ Less than two tables found – caption cannot be correct")
        return score  # cannot meet other requirements anyway

    # ---------- Step 2: Locate blocks & second table index ----------
    blocks = _extract_document_blocks(file_path)
    second_tbl_idx = [i for i, (k, _) in enumerate(blocks) if k == "tbl"][1]
    print(f"✓ Second table located at block index {second_tbl_idx}")

    # ---------- Step 3: Find caption paragraph after second table ----------
    caption_regex = re.compile(r"Table\s*2\s*[:\-–]\s*Metrics", re.IGNORECASE)
    caption_idx, caption_text = None, None
    for idx in range(second_tbl_idx + 1, len(blocks)):
        kind, text = blocks[idx]
        if kind != "p":
            continue
        if text.strip() == "":  # blank paragraph – may occur between table & caption
            continue
        if caption_regex.search(text):
            caption_idx, caption_text = idx, text.strip()
            break
        # Encountered a non-blank paragraph that isn't the caption → stop searching
        break

    if caption_idx is None:
        print("✗ Required caption text 'Table 2: Metrics' not found below second table")
        return score

    print(f"✓ Found caption paragraph at block index {caption_idx}: '{caption_text}'")
    score += 0.3  # correct caption text

    # ---------- Step 4: Verify caption paragraph style ----------
    caption_style_ok = False
    for p in doc.paragraphs:
        if caption_regex.search(p.text):
            style_name = getattr(p.style, "name", "") if p.style else ""
            if "caption" in style_name.lower():
                caption_style_ok = True
                print(f"✓ Caption paragraph uses style '{style_name}'")
                break
            else:
                print(
                    f"✗ Caption paragraph does not use a Caption style (found '{style_name}')"
                )
                break

    if caption_style_ok:
        score += 0.3
    else:
        # Style is important but allow partial credit for correct text/position
        return min(score, max_score)

    # ---------- Step 5: Verify position (immediately below second table) ----------
    immediate_ok = False
    if caption_idx == second_tbl_idx + 1:
        immediate_ok = True
    elif (
        caption_idx == second_tbl_idx + 2
        and blocks[second_tbl_idx + 1][0] == "p"
        and blocks[second_tbl_idx + 1][1].strip() == ""
    ):
        # one blank paragraph allowed between table and caption
        immediate_ok = True

    if immediate_ok:
        print("✓ Caption appears immediately below the second table")
        score += 0.3
    else:
        print("✗ Caption is not positioned directly below the second table")

    final_score = min(score, max_score)
    print(f"REWARD: {final_score}")
    return final_score


# -------------- Run verification when script is executed --------------
if __name__ == "__main__":
    FILE_PATH = (
        "/home/user/libreoffice_writer_question_i_need_the_line_table_2_metrics_to_show_up_as_an_official_caption_right_.docx"
    )
    verify_table_caption_task(FILE_PATH)

