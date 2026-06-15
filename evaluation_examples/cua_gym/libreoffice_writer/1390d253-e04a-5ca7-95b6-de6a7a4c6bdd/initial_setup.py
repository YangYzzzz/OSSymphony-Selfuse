"""
Initial Setup: Create a Writer document with article text (no Pull Quote style).
Task ID: writer_bs_090
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_090'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading('The Future of Sustainable Architecture', level=1)

    # Author byline
    byline = doc.add_paragraph()
    byline.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = byline.add_run('By Elena Vasquez | March 15, 2025')
    run.italic = True
    run.font.size = Pt(10)

    # Article body
    doc.add_paragraph(
        'As cities around the world grapple with the dual challenges of rapid urbanization '
        'and climate change, architects and urban planners are increasingly turning to '
        'sustainable design principles that promise to reshape our built environment. From '
        'living walls in Singapore to passive house standards sweeping across Northern Europe, '
        'the movement toward green architecture has gained unprecedented momentum in recent years.'
    )

    doc.add_paragraph(
        'The concept of biophilic design, which seeks to integrate natural elements into the '
        'built environment, has emerged as one of the most compelling approaches. Research '
        'conducted by the University of Melbourne found that workers in offices with natural '
        'light and greenery reported a 15% increase in productivity and a 23% reduction in '
        'stress-related symptoms compared to those in conventional office spaces.'
    )

    doc.add_heading('Innovative Materials Leading the Way', level=2)

    doc.add_paragraph(
        'Cross-laminated timber, or CLT, has rapidly become the material of choice for '
        'architects committed to reducing the carbon footprint of new construction. Unlike '
        'traditional concrete and steel, which account for approximately 8% of global CO2 '
        'emissions, engineered wood products actually sequester carbon throughout their lifespan. '
        'The Mjostaarnet tower in Brumunddal, Norway, standing at 85.4 meters, demonstrates '
        'that timber construction can reach impressive heights while maintaining structural integrity.'
    )

    doc.add_paragraph(
        'Self-healing concrete represents another breakthrough in sustainable building materials. '
        'Developed by researchers at Delft University of Technology, this innovative material '
        'contains limestone-producing bacteria that activate when cracks form, automatically '
        'sealing gaps and extending the lifespan of structures by decades. Early field trials '
        'in the Netherlands have shown a 60% reduction in maintenance costs over a ten-year period.'
    )

    doc.add_heading('Energy Efficiency and Net-Zero Buildings', level=2)

    doc.add_paragraph(
        'The push toward net-zero energy buildings has accelerated dramatically since the Paris '
        'Agreement. In 2024, the International Energy Agency reported that buildings certified '
        'under green building standards consumed 40% less energy than their conventional '
        'counterparts. Countries like Denmark and Germany have already mandated that all new '
        'public buildings achieve net-zero energy status by 2030.'
    )

    doc.add_paragraph(
        'Passive design strategies, including optimal building orientation, high-performance '
        'insulation, and advanced glazing systems, form the foundation of energy-efficient '
        'architecture. The Bullitt Center in Seattle, often called the greenest commercial '
        'building in the world, generates more energy than it consumes through a combination '
        'of rooftop solar panels and aggressive energy conservation measures.'
    )

    doc.add_heading('Community-Centered Design', level=2)

    doc.add_paragraph(
        'Perhaps the most significant shift in sustainable architecture is the growing emphasis '
        'on community engagement and social equity. Architects like Diebedo Francis Kere, '
        'winner of the 2022 Pritzker Architecture Prize, have demonstrated that sustainable '
        'design need not be the exclusive domain of wealthy nations. His projects in Burkina '
        'Faso use locally sourced materials and community labor to create schools, health '
        'centers, and public gathering spaces that are both environmentally responsible and '
        'culturally meaningful.'
    )

    doc.add_paragraph(
        'As we look toward the next decade, the integration of artificial intelligence and '
        'digital twin technology promises to further optimize building performance. Smart '
        'buildings equipped with sensor networks and machine learning algorithms can '
        'continuously adapt their energy use, ventilation, and lighting to occupant needs, '
        'reducing waste while improving comfort. The future of architecture, it seems, lies '
        'not just in the materials we use, but in how intelligently we deploy them.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
