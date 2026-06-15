"""
Reward Script: Set body text font to Calibri 11pt in offer letter
Task ID: writer_hr_001
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Normal style font name changed to Calibri
  Component 2 (0.3): Normal style font size changed to 11pt
  Component 3 (0.2): All body-text runs use Calibri font
  Component 4 (0.2): All body-text runs use 11pt size
"""

import os

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_001'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that body text font has been changed to Calibri 11pt.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Normal style font name is Calibri (0.3 points)
    # Initial: Liberation Sans -> Golden: Calibri
    try:
        normal_style = doc.styles['Normal']
        style_font_name = normal_style.font.name
        if style_font_name and style_font_name.lower() == 'calibri':
            print(f"PASS: Component 1 — Normal style font name is '{style_font_name}' (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Expected Normal style font 'Calibri', found: '{style_font_name}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Normal style font size is 11pt (0.3 points)
    # Initial: 10pt -> Golden: 11pt
    try:
        normal_style = doc.styles['Normal']
        style_font_size = normal_style.font.size
        if style_font_size and abs(style_font_size.pt - 11.0) < 0.1:
            print(f"PASS: Component 2 — Normal style font size is {style_font_size.pt}pt (0.3 pts)")
            total_score += 0.3
        else:
            size_str = f"{style_font_size.pt}pt" if style_font_size else "None"
            print(f"FAIL: Component 2 — Expected Normal style size 11pt, found: {size_str}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Effective font of body-text runs is Calibri (0.2 points)
    # "Body text" = runs that are NOT the 16pt title or 9pt address header.
    # We resolve the effective font: explicit run font > style font.
    # Initial: effective font is Liberation Sans. Golden: effective font is Calibri.
    try:
        normal_font_name = doc.styles['Normal'].font.name  # fallback for inherited
        body_runs_total = 0
        body_runs_calibri = 0
        for para in doc.paragraphs:
            for run in para.runs:
                # Resolve effective font name
                effective_name = run.font.name if run.font.name else normal_font_name
                # Resolve effective size to identify body text
                effective_size = run.font.size.pt if run.font.size else (
                    doc.styles['Normal'].font.size.pt if doc.styles['Normal'].font.size else None
                )
                # Skip non-body runs (title ~16pt, address ~9pt)
                if effective_size and (effective_size > 14.0 or effective_size < 9.5):
                    continue
                body_runs_total += 1
                if effective_name and effective_name.lower() == 'calibri':
                    body_runs_calibri += 1

        if body_runs_total > 0 and body_runs_calibri == body_runs_total:
            print(f"PASS: Component 3 — All {body_runs_total} body-text runs effectively use Calibri (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — {body_runs_calibri}/{body_runs_total} body-text runs effectively use Calibri")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Effective font size of body-text runs is 11pt (0.2 points)
    # Initial body text was effectively 10pt, golden should be 11pt.
    # Resolve effective size: explicit run size > style size.
    try:
        normal_font_size_pt = doc.styles['Normal'].font.size.pt if doc.styles['Normal'].font.size else None
        body_runs_total = 0
        body_runs_11pt = 0
        for para in doc.paragraphs:
            for run in para.runs:
                effective_size = run.font.size.pt if run.font.size else normal_font_size_pt
                # Skip non-body runs (title ~16pt, address ~9pt)
                if effective_size and (effective_size > 14.0 or effective_size < 9.5):
                    continue
                body_runs_total += 1
                if effective_size and abs(effective_size - 11.0) < 0.1:
                    body_runs_11pt += 1

        if body_runs_total > 0 and body_runs_11pt == body_runs_total:
            print(f"PASS: Component 4 — All {body_runs_total} body-text runs effectively 11pt (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 4 — {body_runs_11pt}/{body_runs_total} body-text runs at 11pt")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    # Round to avoid floating point artifacts
    final_score = round(final_score, 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
