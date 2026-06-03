"""
Initial Setup: Offer letter with body text in Liberation Sans 10pt
Task ID: writer_hr_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

BODY_FONT = 'Liberation Sans'
BODY_SIZE = Pt(10)


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_paragraph(doc, text, bold=False, alignment=None, space_after=Pt(6)):
    """Add a paragraph with Liberation Sans 10pt formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.bold = bold
    para.paragraph_format.space_after = space_after
    if alignment:
        para.paragraph_format.alignment = alignment
    return para


def create_initial():
    doc = Document()

    # Set default style to Liberation Sans 10pt
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE

    # --- Company Header ---
    header_para = doc.add_paragraph()
    header_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.paragraph_format.space_after = Pt(2)
    run = header_para.add_run('Meridian Technology Solutions')
    run.font.name = BODY_FONT
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    addr_para = doc.add_paragraph()
    addr_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr_para.paragraph_format.space_after = Pt(0)
    run = addr_para.add_run('2400 Innovation Drive, Suite 800')
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    addr2_para = doc.add_paragraph()
    addr2_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr2_para.paragraph_format.space_after = Pt(12)
    run = addr2_para.add_run('San Francisco, CA 94105  |  (415) 555-0192  |  hr@meridiantech.com')
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # --- Date ---
    add_body_paragraph(doc, 'March 28, 2026', space_after=Pt(12))

    # --- Recipient ---
    add_body_paragraph(doc, 'Ms. Elena Rodriguez', space_after=Pt(0))
    add_body_paragraph(doc, '1847 Oakwood Lane, Apt 3B', space_after=Pt(0))
    add_body_paragraph(doc, 'Portland, OR 97205', space_after=Pt(12))

    # --- Subject ---
    add_body_paragraph(doc, 'RE: Offer of Employment — Senior Software Engineer', bold=True,
                       space_after=Pt(12))

    # --- Greeting ---
    add_body_paragraph(doc, 'Dear Ms. Rodriguez,', space_after=Pt(6))

    # --- Body Paragraphs ---
    add_body_paragraph(doc,
        'On behalf of Meridian Technology Solutions, I am pleased to extend this offer of '
        'employment for the position of Senior Software Engineer within our Platform Engineering '
        'division. We were very impressed with your technical expertise and collaborative approach '
        'during the interview process, and we believe you will be an excellent addition to our team.')

    add_body_paragraph(doc,
        'This letter outlines the key terms and conditions of your employment with Meridian '
        'Technology Solutions. Please review the details carefully and feel free to reach out '
        'with any questions.')

    # --- Position Details ---
    add_body_paragraph(doc, 'Position Details', bold=True, space_after=Pt(4))

    details = [
        ('Title:', 'Senior Software Engineer'),
        ('Department:', 'Platform Engineering'),
        ('Reports to:', 'Dr. James Whitfield, VP of Engineering'),
        ('Start Date:', 'April 21, 2026'),
        ('Location:', 'San Francisco, CA (hybrid — 3 days in-office)'),
    ]
    for label, value in details:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Inches(0.5)
        run_label = para.add_run(label + '  ')
        run_label.font.name = BODY_FONT
        run_label.font.size = BODY_SIZE
        run_label.bold = True
        run_val = para.add_run(value)
        run_val.font.name = BODY_FONT
        run_val.font.size = BODY_SIZE

    doc.add_paragraph()  # spacer

    # --- Compensation ---
    add_body_paragraph(doc, 'Compensation and Benefits', bold=True, space_after=Pt(4))

    add_body_paragraph(doc,
        'Your annual base salary will be $165,000.00, paid on a bi-weekly basis. In addition, '
        'you will be eligible for an annual performance bonus of up to 15% of your base salary, '
        'subject to individual and company performance targets.')

    add_body_paragraph(doc,
        'You will also receive an equity grant of 12,000 stock options, vesting over a four-year '
        'period with a one-year cliff. The grant is subject to the terms of the Meridian Technology '
        'Solutions 2024 Equity Incentive Plan and approval by the Board of Directors.')

    add_body_paragraph(doc,
        'Our comprehensive benefits package includes medical, dental, and vision coverage '
        'effective on your first day of employment; 401(k) with 4% company match; 20 days of '
        'paid time off annually; 10 paid holidays; and a $2,500 annual professional development '
        'stipend. Full details are provided in the enclosed Benefits Summary.')

    # --- Conditions ---
    add_body_paragraph(doc, 'Conditions of Employment', bold=True, space_after=Pt(4))

    add_body_paragraph(doc,
        'This offer is contingent upon the successful completion of a standard background '
        'verification and reference check. Employment at Meridian Technology Solutions is '
        'at-will, meaning either party may terminate the relationship at any time with or '
        'without cause or notice.')

    add_body_paragraph(doc,
        'On your first day, you will be required to complete an I-9 Employment Eligibility '
        'Verification form and provide acceptable documentation establishing your identity '
        'and work authorization.')

    # --- Acceptance ---
    add_body_paragraph(doc, 'Acceptance', bold=True, space_after=Pt(4))

    add_body_paragraph(doc,
        'To accept this offer, please sign and date this letter below and return it to our '
        'Human Resources department no later than April 7, 2026. If we do not receive your '
        'signed acceptance by this date, the offer will be considered withdrawn.')

    add_body_paragraph(doc,
        'We are excited about the possibility of you joining Meridian Technology Solutions and '
        'look forward to your positive response.', space_after=Pt(12))

    # --- Closing ---
    add_body_paragraph(doc, 'Sincerely,', space_after=Pt(24))

    add_body_paragraph(doc, 'Catherine Park', bold=True, space_after=Pt(0))
    add_body_paragraph(doc, 'Director of Human Resources', space_after=Pt(0))
    add_body_paragraph(doc, 'Meridian Technology Solutions', space_after=Pt(24))

    # --- Acceptance Block ---
    add_body_paragraph(doc, 'ACCEPTED AND AGREED:', bold=True, space_after=Pt(18))

    add_body_paragraph(doc, '______________________________', space_after=Pt(0))
    add_body_paragraph(doc, 'Elena Rodriguez', space_after=Pt(12))

    add_body_paragraph(doc, 'Date: ______________________________', space_after=Pt(0))

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
