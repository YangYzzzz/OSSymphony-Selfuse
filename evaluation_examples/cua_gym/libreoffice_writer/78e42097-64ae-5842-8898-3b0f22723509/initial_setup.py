"""
Initial Setup: Create a Writer document with product-specific terms that will be flagged by spell check.
Task ID: writer_tech_094
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_094'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Nextera Platform — Product Terminology Guide', level=1)

    # Introduction paragraph
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        'This document provides an overview of key product names and technologies '
        'within the Nextera cloud infrastructure platform. All team members should '
        'familiarize themselves with these terms for consistent usage in documentation, '
        'presentations, and customer-facing communications.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Section 1: CloudSync
    doc.add_heading('1. CloudSync', level=2)
    p1 = doc.add_paragraph()
    r1 = p1.add_run(
        'CloudSync is our flagship real-time data synchronization engine. It enables '
        'seamless replication of datasets across multiple cloud regions with sub-second '
        'latency. The CloudSync service handles conflict resolution automatically using '
        'a last-writer-wins strategy combined with vector clocks for causal ordering.'
    )
    r1.font.size = Pt(11)
    r1.font.name = 'Calibri'

    p1b = doc.add_paragraph()
    r1b = p1b.add_run(
        'Deployment note: CloudSync requires at least three availability zones for '
        'production workloads. Contact the infrastructure team before provisioning a '
        'new CloudSync cluster.'
    )
    r1b.font.size = Pt(11)
    r1b.font.name = 'Calibri'

    # Section 2: DataMesh
    doc.add_heading('2. DataMesh', level=2)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'DataMesh is our distributed data governance layer that implements domain-oriented '
        'data ownership. Unlike traditional centralized data lakes, DataMesh treats data '
        'as a product, with each domain team responsible for publishing high-quality, '
        'discoverable datasets through standardized APIs.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'

    p2b = doc.add_paragraph()
    r2b = p2b.add_run(
        'The DataMesh architecture integrates with our existing ETL pipelines and supports '
        'both batch and streaming ingestion patterns. Teams adopting DataMesh should review '
        'the data product specification template in Confluence.'
    )
    r2b.font.size = Pt(11)
    r2b.font.name = 'Calibri'

    # Section 3: AutoScale
    doc.add_heading('3. AutoScale', level=2)
    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        'AutoScale is the intelligent resource management module that dynamically adjusts '
        'compute capacity based on real-time demand signals. AutoScale uses predictive '
        'machine learning models trained on historical traffic patterns to pre-provision '
        'resources before demand spikes occur.'
    )
    r3.font.size = Pt(11)
    r3.font.name = 'Calibri'

    p3b = doc.add_paragraph()
    r3b = p3b.add_run(
        'Key metrics monitored by AutoScale include CPU utilization, memory pressure, '
        'request queue depth, and custom application-level indicators. The AutoScale '
        'dashboard is accessible via the platform control panel under Infrastructure > Scaling.'
    )
    r3b.font.size = Pt(11)
    r3b.font.name = 'Calibri'

    # Section 4: NetGuard
    doc.add_heading('4. NetGuard', level=2)
    p4 = doc.add_paragraph()
    r4 = p4.add_run(
        'NetGuard is our zero-trust network security framework that enforces microsegmentation '
        'and continuous authentication across all service-to-service communications. NetGuard '
        'implements mutual TLS by default and integrates with the identity provider for '
        'fine-grained access control policies.'
    )
    r4.font.size = Pt(11)
    r4.font.name = 'Calibri'

    p4b = doc.add_paragraph()
    r4b = p4b.add_run(
        'All production services must be enrolled in NetGuard before their first deployment. '
        'The NetGuard agent runs as a sidecar container and requires no application code changes. '
        'Security audits can be triggered through the NetGuard compliance portal.'
    )
    r4b.font.size = Pt(11)
    r4b.font.name = 'Calibri'

    # Section 5: LogStream
    doc.add_heading('5. LogStream', level=2)
    p5 = doc.add_paragraph()
    r5 = p5.add_run(
        'LogStream is the centralized log aggregation and analysis platform. It collects, '
        'indexes, and stores structured logs from all microservices in the Nextera ecosystem. '
        'LogStream supports real-time alerting, full-text search, and custom dashboards for '
        'operational monitoring.'
    )
    r5.font.size = Pt(11)
    r5.font.name = 'Calibri'

    p5b = doc.add_paragraph()
    r5b = p5b.add_run(
        'Teams should configure their services to emit logs in the LogStream JSON format. '
        'The LogStream SDK provides language-specific wrappers for Java, Python, Go, and '
        'TypeScript. Retention policies are managed at the namespace level — contact the '
        'platform team to adjust retention for your LogStream namespace.'
    )
    r5b.font.size = Pt(11)
    r5b.font.name = 'Calibri'

    # Summary section
    doc.add_heading('Summary', level=2)
    summary = doc.add_paragraph()
    rs = summary.add_run(
        'The five core platform components — CloudSync, DataMesh, AutoScale, NetGuard, '
        'and LogStream — form the backbone of the Nextera infrastructure. Each component '
        'is designed to operate independently while providing deep integration points with '
        'the others. Refer to the individual product documentation for API references and '
        'configuration guides.'
    )
    rs.font.size = Pt(11)
    rs.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure no custom dictionary exists with these terms
    # Remove any existing custom dictionaries that might have them
    dic_dir = os.path.expanduser('~/.config/libreoffice/4/user/wordbook')
    if os.path.isdir(dic_dir):
        for fname in os.listdir(dic_dir):
            fpath = os.path.join(dic_dir, fname)
            if fname.endswith('.dic') and fname != 'standard.dic':
                # Check if it contains our terms
                try:
                    with open(fpath, 'r', errors='ignore') as f:
                        content = f.read()
                    if any(term in content for term in ['CloudSync', 'DataMesh', 'AutoScale', 'NetGuard', 'LogStream']):
                        os.remove(fpath)
                        print(f'Removed existing dictionary: {fpath}')
                except Exception:
                    pass

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
