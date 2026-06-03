"""
Reward Script: Workshop handout — 'Introduction to Public Speaking'
Task ID: writer_wf_082
Domain: libreoffice_writer
Scoring:
  C1: Title as Heading 1 (0.10)
  C2: Presenter name and date (0.05)
  C3: 6 Heading 2 sections (0.15)
  C4: Key Principles — 5 numbered items (0.15)
  C5: Common Mistakes — 4 bulleted items (0.10)
  C6: Two exercise sections with instructions (0.10)
  C7: Tips for Managing Nervousness — 6 bulleted tips (0.10)
  C8: Recommended Resources — 4 bulleted items (0.10)
  C9: Self-assessment table with 3 cols and 5 skill rows (0.15)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_082'


def persist_app_state(domain: str):
    """Best-effort save in case LibreOffice has unsaved changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build helper data structures
    paragraphs = doc.paragraphs
    styles = [(p.style.name if p.style else 'Normal', p.text) for p in paragraphs]

    heading2_sections = []
    current_h2 = None
    current_items = []
    for style_name, text in styles:
        if style_name == 'Heading 2':
            if current_h2 is not None:
                heading2_sections.append((current_h2, current_items))
            current_h2 = text
            current_items = []
        elif current_h2 is not None:
            current_items.append((style_name, text))
    if current_h2 is not None:
        heading2_sections.append((current_h2, current_items))

    h2_titles = [h2 for h2, _ in heading2_sections]
    h2_map = {h2.lower().strip(): items for h2, items in heading2_sections}

    # ============================================================
    # Component 1: Title as Heading 1 (0.10 points)
    # ============================================================
    try:
        title_matches = [t for s, t in styles if s.startswith('Heading 1') and 'public speaking' in t.lower()]
        if len(title_matches) > 0:
            print(f"PASS: Component 1 — Title 'Introduction to Public Speaking' as Heading 1 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — No Heading 1 containing 'public speaking' found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ============================================================
    # Component 2: Presenter name and date present (0.05 points)
    # ============================================================
    try:
        # Look for any paragraph (before first Heading 2) containing presenter-like
        # info and a date-like string
        preamble_texts = []
        for style_name, text in styles:
            if style_name == 'Heading 2':
                break
            if style_name != 'Heading 1' and text.strip():
                preamble_texts.append(text.lower())

        all_preamble = ' '.join(preamble_texts)
        has_presenter = any(kw in all_preamble for kw in ['presenter', 'instructor', 'facilitator', 'name', 'dr.', 'prof.'])
        has_date = any(kw in all_preamble for kw in ['date', '2025', '2026', 'january', 'february', 'march', 'april', 'may', 'june',
                                                       'july', 'august', 'september', 'october', 'november', 'december'])

        if has_presenter and has_date:
            print(f"PASS: Component 2 — Presenter and date found in preamble (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 2 — Presenter={has_presenter}, Date={has_date} in preamble: {all_preamble[:100]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ============================================================
    # Component 3: 6 Heading 2 sections (0.15 points)
    # ============================================================
    try:
        num_h2 = len(h2_titles)
        if num_h2 >= 6:
            print(f"PASS: Component 3 — Found {num_h2} Heading 2 sections (>= 6) (0.15 pts)")
            total_score += 0.15
        elif num_h2 >= 4:
            partial = round(0.15 * num_h2 / 6, 2)
            print(f"PARTIAL: Component 3 — Found {num_h2}/6 Heading 2 sections ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Found only {num_h2} Heading 2 sections, need >= 6")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ============================================================
    # Component 4: Key Principles — 5 numbered items (0.15 points)
    # ============================================================
    try:
        kp_items = None
        for h2_title, items in heading2_sections:
            if 'key principles' in h2_title.lower() or 'principles' in h2_title.lower():
                kp_items = items
                break

        if kp_items is not None:
            numbered = [t for s, t in kp_items if 'list number' in s.lower() and t.strip()]
            if len(numbered) >= 5:
                print(f"PASS: Component 4 — Key Principles has {len(numbered)} numbered items (>= 5) (0.15 pts)")
                total_score += 0.15
            elif len(numbered) >= 3:
                partial = round(0.15 * len(numbered) / 5, 2)
                print(f"PARTIAL: Component 4 — Key Principles has {len(numbered)}/5 numbered items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Key Principles has {len(numbered)} numbered items, need >= 5")
        else:
            print(f"FAIL: Component 4 — No 'Key Principles' section found among: {h2_titles}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ============================================================
    # Component 5: Common Mistakes — 4 bulleted items (0.10 points)
    # ============================================================
    try:
        cm_items = None
        for h2_title, items in heading2_sections:
            if 'common mistakes' in h2_title.lower() or 'mistakes' in h2_title.lower():
                cm_items = items
                break

        if cm_items is not None:
            bulleted = [t for s, t in cm_items if 'list bullet' in s.lower() and t.strip()]
            if len(bulleted) >= 4:
                print(f"PASS: Component 5 — Common Mistakes has {len(bulleted)} bulleted items (>= 4) (0.10 pts)")
                total_score += 0.10
            elif len(bulleted) >= 2:
                partial = round(0.10 * len(bulleted) / 4, 2)
                print(f"PARTIAL: Component 5 — Common Mistakes has {len(bulleted)}/4 bulleted items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — Common Mistakes has {len(bulleted)} bulleted items, need >= 4")
        else:
            print(f"FAIL: Component 5 — No 'Common Mistakes' section found among: {h2_titles}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ============================================================
    # Component 6: Two exercise sections with instructions (0.10 points)
    # ============================================================
    try:
        exercise_count = 0
        exercise_has_instructions = 0
        for h2_title, items in heading2_sections:
            if 'exercise' in h2_title.lower():
                exercise_count += 1
                # Check if there's at least one Normal paragraph with instruction text
                normal_texts = [t for s, t in items if t.strip() and 'heading' not in s.lower()]
                if any('instruct' in t.lower() or len(t) > 20 for t in normal_texts):
                    exercise_has_instructions += 1

        if exercise_count >= 2 and exercise_has_instructions >= 2:
            print(f"PASS: Component 6 — {exercise_count} exercise sections, {exercise_has_instructions} with instructions (0.10 pts)")
            total_score += 0.10
        elif exercise_count >= 1:
            partial = round(0.10 * min(exercise_has_instructions, 2) / 2, 2)
            print(f"PARTIAL: Component 6 — {exercise_count} exercises, {exercise_has_instructions} with instructions ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 — Found {exercise_count} exercise sections, need >= 2")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # ============================================================
    # Component 7: Tips for Managing Nervousness — 6 bulleted tips (0.10 points)
    # ============================================================
    try:
        tips_items = None
        for h2_title, items in heading2_sections:
            if 'nervousness' in h2_title.lower() or 'managing' in h2_title.lower():
                tips_items = items
                break

        if tips_items is not None:
            bulleted = [t for s, t in tips_items if 'list bullet' in s.lower() and t.strip()]
            if len(bulleted) >= 6:
                print(f"PASS: Component 7 — Tips has {len(bulleted)} bulleted items (>= 6) (0.10 pts)")
                total_score += 0.10
            elif len(bulleted) >= 3:
                partial = round(0.10 * len(bulleted) / 6, 2)
                print(f"PARTIAL: Component 7 — Tips has {len(bulleted)}/6 bulleted items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 — Tips has {len(bulleted)} bulleted items, need >= 6")
        else:
            print(f"FAIL: Component 7 — No 'Tips for Managing Nervousness' section found among: {h2_titles}")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # ============================================================
    # Component 8: Recommended Resources — 4 bulleted items (0.10 points)
    # ============================================================
    try:
        res_items = None
        for h2_title, items in heading2_sections:
            if 'resource' in h2_title.lower() or 'recommended' in h2_title.lower():
                res_items = items
                break

        if res_items is not None:
            bulleted = [t for s, t in res_items if 'list bullet' in s.lower() and t.strip()]
            if len(bulleted) >= 4:
                print(f"PASS: Component 8 — Resources has {len(bulleted)} bulleted items (>= 4) (0.10 pts)")
                total_score += 0.10
            elif len(bulleted) >= 2:
                partial = round(0.10 * len(bulleted) / 4, 2)
                print(f"PARTIAL: Component 8 — Resources has {len(bulleted)}/4 bulleted items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 8 — Resources has {len(bulleted)} bulleted items, need >= 4")
        else:
            print(f"FAIL: Component 8 — No 'Recommended Resources' section found among: {h2_titles}")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    # ============================================================
    # Component 9: Self-assessment table — 3 cols, header + 5 skill rows (0.15 points)
    # ============================================================
    try:
        tables = doc.tables
        if len(tables) >= 1:
            # Find the self-assessment table (any table with "Skill" header and rating columns)
            sa_table = None
            for table in tables:
                header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'skill' in header_cells and any('rating' in c or 'before' in c for c in header_cells):
                    sa_table = table
                    break

            if sa_table is None:
                # Fallback: use the last table if it has 3 columns
                for table in tables:
                    if len(table.columns) == 3 and len(table.rows) >= 2:
                        sa_table = table
                        break

            if sa_table is not None:
                num_cols = len(sa_table.columns)
                num_rows = len(sa_table.rows)
                # Expect header + 5 skill rows = 6 rows, 3 cols
                col_ok = (num_cols == 3)
                rows_ok = (num_rows >= 6)  # header + 5 skills
                # Check that skill rows have non-empty first column
                skill_rows = 0
                for ri in range(1, num_rows):
                    cell_text = sa_table.cell(ri, 0).text.strip()
                    if cell_text:
                        skill_rows += 1

                if col_ok and skill_rows >= 5:
                    print(f"PASS: Component 9 — Self-assessment table: {num_cols} cols, {skill_rows} skill rows (0.15 pts)")
                    total_score += 0.15
                elif col_ok and skill_rows >= 3:
                    partial = round(0.15 * skill_rows / 5, 2)
                    print(f"PARTIAL: Component 9 — Table: {num_cols} cols, {skill_rows}/5 skill rows ({partial} pts)")
                    total_score += partial
                else:
                    print(f"FAIL: Component 9 — Table has {num_cols} cols (need 3), {skill_rows} skill rows (need 5)")
            else:
                print(f"FAIL: Component 9 — No self-assessment table found")
        else:
            print(f"FAIL: Component 9 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 9 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
