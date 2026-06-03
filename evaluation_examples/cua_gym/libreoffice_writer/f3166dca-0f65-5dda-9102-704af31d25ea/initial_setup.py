"""
Initial Setup: Insert page count field in footer
Task ID: writer_acad_067
Domain: libreoffice_writer

Creates a thesis document with a footer containing only the current page number.
The agent must modify the footer to read 'Page X of Y'.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_067'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph (current page number only)."""
    # begin
    r1 = paragraph.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)
    # instruction
    r2 = paragraph.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)
    # end
    r3 = paragraph.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Footer with page number only ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp)

    # --- Title Page ---
    title = doc.add_heading('Adaptive Neural Network Architectures for Real-Time Climate Prediction Models', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Thesis Submitted in Partial Fulfillment of the Requirements\n'
                           'for the Degree of Doctor of Philosophy\n\n'
                           'Department of Computer Science\n'
                           'Stanford University\n\n'
                           'Elena Vasquez\n'
                           'March 2026')
    run.font.size = Pt(12)

    # --- Abstract ---
    doc.add_page_break()
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'Climate prediction remains one of the most computationally demanding challenges in '
        'modern science. Traditional numerical weather prediction (NWP) models require enormous '
        'computational resources and often fail to capture fine-grained regional patterns. This '
        'dissertation presents a novel adaptive neural network framework, ClimateNet-A, that '
        'dynamically adjusts its architecture based on atmospheric complexity metrics. Our approach '
        'reduces inference time by 73% compared to standard deep learning baselines while maintaining '
        'prediction accuracy within 0.4°C RMSE for 72-hour forecasts across 14 global climate zones.'
    )
    doc.add_paragraph(
        'We introduce three key innovations: (1) a hierarchical attention mechanism that captures '
        'multi-scale atmospheric interactions, (2) a physics-informed loss function incorporating '
        'conservation laws from thermodynamics and fluid dynamics, and (3) an adaptive computation '
        'module that allocates network capacity proportionally to local atmospheric instability. '
        'Experiments on ERA5 reanalysis data (1979–2024) demonstrate that ClimateNet-A achieves '
        'state-of-the-art performance on precipitation forecasting (CSI-0.5 = 0.847) and extreme '
        'weather event detection (F1 = 0.912).'
    )

    # --- Chapter 1: Introduction ---
    doc.add_page_break()
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_heading('1.1 Motivation', level=2)
    doc.add_paragraph(
        'The global climate system is a complex, nonlinear dynamical system whose behavior emerges '
        'from interactions across spatial scales spanning seven orders of magnitude—from cloud '
        'microphysics at the millimeter scale to planetary wave patterns spanning thousands of '
        'kilometers. Accurate climate prediction is critical for agriculture, disaster preparedness, '
        'energy infrastructure planning, and public health interventions.'
    )
    doc.add_paragraph(
        'Current operational forecasting systems, such as the European Centre for Medium-Range '
        'Weather Forecasts (ECMWF) Integrated Forecasting System (IFS) and the National Centers '
        'for Environmental Prediction (NCEP) Global Forecast System (GFS), solve partial differential '
        'equations governing atmospheric motion on discretized grids. While these systems have '
        'steadily improved over decades, they face fundamental limitations in computational '
        'scalability and the representation of sub-grid-scale processes through parameterization '
        'schemes (Bauer et al., 2015).'
    )

    doc.add_heading('1.2 Research Questions', level=2)
    doc.add_paragraph(
        'This thesis addresses the following research questions:'
    )
    doc.add_paragraph('RQ1: Can adaptive neural architectures match or exceed the accuracy of '
                      'fixed-topology deep learning models for multi-variable climate prediction?',
                      style='List Number')
    doc.add_paragraph('RQ2: How can physical conservation laws be effectively integrated into '
                      'neural network training objectives without sacrificing gradient flow stability?',
                      style='List Number')
    doc.add_paragraph('RQ3: What is the relationship between local atmospheric complexity and '
                      'optimal network capacity allocation?',
                      style='List Number')

    doc.add_heading('1.3 Contributions', level=2)
    doc.add_paragraph(
        'The principal contributions of this dissertation are threefold. First, we propose '
        'ClimateNet-A, an adaptive architecture that dynamically adjusts layer depth and attention '
        'head count based on a learned complexity estimator. Second, we formulate a physics-informed '
        'loss landscape that enforces mass, energy, and momentum conservation without requiring '
        'explicit PDE solvers during training. Third, we provide a comprehensive empirical analysis '
        'across 14 climate zones demonstrating consistent improvements in forecast skill metrics.'
    )

    # --- Chapter 2: Literature Review ---
    doc.add_page_break()
    doc.add_heading('Chapter 2: Literature Review', level=1)
    doc.add_heading('2.1 Numerical Weather Prediction', level=2)
    doc.add_paragraph(
        'The foundations of numerical weather prediction were laid by Vilhelm Bjerknes in 1904, '
        'who formulated the primitive equations governing atmospheric motion. Lewis Fry Richardson '
        'attempted the first numerical forecast in 1922, computing by hand a six-hour pressure '
        'change that proved wildly inaccurate due to initialization errors. It was not until 1950 '
        'that Charney, Fjørtoft, and von Neumann produced the first successful computer-based '
        'weather forecast using the ENIAC computer at the Aberdeen Proving Ground.'
    )
    doc.add_paragraph(
        'Modern NWP systems discretize the atmosphere into three-dimensional grid cells and '
        'integrate the governing equations forward in time. The ECMWF IFS currently operates at '
        'approximately 9 km horizontal resolution with 137 vertical levels, requiring approximately '
        '4 million core-hours per 10-day forecast on Atos supercomputer systems (ECMWF, 2023). '
        'Despite these resources, systematic biases persist in tropical convection, polar vortex '
        'dynamics, and boundary layer turbulence representation.'
    )

    doc.add_heading('2.2 Deep Learning for Weather Prediction', level=2)
    doc.add_paragraph(
        'The application of deep learning to weather prediction has accelerated dramatically since '
        '2020. Pathak et al. (2022) demonstrated that a FourCastNet model based on Fourier neural '
        'operators could produce global forecasts at 0.25° resolution in under two seconds, compared '
        'to several hours for NWP. Lam et al. (2023) introduced GraphCast, a graph neural network '
        'that achieved ten-day forecast accuracy surpassing ECMWF HRES on 90% of atmospheric '
        'variables. Bi et al. (2023) proposed Pangu-Weather, a 3D Swin Transformer achieving '
        'similar results with a different architectural approach.'
    )

    # --- Chapter 3: Methodology ---
    doc.add_page_break()
    doc.add_heading('Chapter 3: Methodology', level=1)
    doc.add_heading('3.1 Architecture Overview', level=2)
    doc.add_paragraph(
        'ClimateNet-A comprises four primary components: (1) a multi-resolution encoder that '
        'processes atmospheric state variables at multiple spatial scales, (2) a complexity estimator '
        'module that computes local atmospheric instability metrics, (3) an adaptive transformer '
        'backbone whose depth and width are modulated by the complexity estimates, and (4) a '
        'multi-variable decoder that produces forecasts for temperature, geopotential, wind '
        'components, and specific humidity at each pressure level.'
    )

    doc.add_heading('3.2 Adaptive Computation Module', level=2)
    doc.add_paragraph(
        'The adaptive computation module is inspired by early-exit networks (Teerapittayanon et al., '
        '2016) and mixture-of-experts architectures (Shazeer et al., 2017). At each transformer '
        'layer, the complexity estimator produces a halting probability σ(x) ∈ [0, 1] for each '
        'spatial location. Computation is allocated dynamically: stable atmospheric regions (e.g., '
        'subtropical high-pressure systems) may require only 4-6 transformer layers, while highly '
        'unstable regions (e.g., developing tropical cyclones) may utilize the full 24-layer stack.'
    )

    doc.add_heading('3.3 Physics-Informed Loss Function', level=2)
    doc.add_paragraph(
        'The total training objective combines a weighted MSE reconstruction loss with three '
        'physics-based penalty terms: mass conservation (enforcing that column-integrated moisture '
        'changes equal precipitation minus evaporation), energy conservation (enforcing that total '
        'moist static energy is approximately conserved along adiabatic trajectories), and vorticity '
        'conservation (enforcing that potential vorticity is conserved in the absence of diabatic '
        'heating and friction). Each penalty is formulated as a soft constraint with a learnable '
        'Lagrange multiplier that adapts during training.'
    )

    # --- References placeholder ---
    doc.add_page_break()
    doc.add_heading('References', level=1)
    doc.add_paragraph(
        'Bauer, P., Thorpe, A., & Brunet, G. (2015). The quiet revolution of numerical weather '
        'prediction. Nature, 525(7567), 47–55.'
    )
    doc.add_paragraph(
        'Bi, K., Xie, L., Zhang, H., et al. (2023). Accurate medium-range global weather '
        'forecasting with 3D neural networks. Nature, 619(7970), 533–538.'
    )
    doc.add_paragraph(
        'Lam, R., Sanchez-Gonzalez, A., Willson, M., et al. (2023). Learning skillful medium-range '
        'global weather forecasting. Science, 382(6677), 1416–1421.'
    )
    doc.add_paragraph(
        'Pathak, J., Subramanian, S., Harrington, P., et al. (2022). FourCastNet: A global '
        'data-driven high-resolution weather forecasting model. arXiv preprint arXiv:2202.11214.'
    )
    doc.add_paragraph(
        'Shazeer, N., Mirhoseini, A., Maczuga, K., et al. (2017). Outrageously large neural '
        'networks: The sparsely-gated mixture-of-experts layer. ICLR 2017.'
    )
    doc.add_paragraph(
        'Teerapittayanon, S., McDanel, B., & Kung, H. T. (2016). BranchyNet: Fast inference '
        'via early exiting from deep neural networks. ICPR 2016.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
