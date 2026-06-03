"""
Reward Script: Add bibliography entries for Kim (2019a, 2019b, 2020) with disambiguation
Task ID: writer_bs_043
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Three Kim bibliography entries present in References section
  Component 2 (0.20): Bibliography entries contain correct titles and years
  Component 3 (0.20): 2019 entries disambiguated with 'a' and 'b' suffixes in bibliography
  Component 4 (0.20): In-text citations (Kim, 2019a; Kim, 2019b; Kim, 2020) in body
  Component 5 (0.10): In-text citation suffixes are consistent with bibliography
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_043'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            import time
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the References section
    ref_start_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'references' and para.style and 'Heading' in para.style.name:
            ref_start_idx = i
            break

    if ref_start_idx is None:
        print("FAIL: Could not locate 'References' heading")
        print("REWARD: 0.0")
        return 0.0

    # Collect reference paragraphs (after References heading)
    ref_paragraphs = []
    for i in range(ref_start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text:
            # Stop if we hit another heading
            style_name = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ''
            if 'Heading' in style_name:
                break
            ref_paragraphs.append(text)

    # Collect body paragraphs (before References heading)
    body_texts = []
    for i in range(0, ref_start_idx):
        text = doc.paragraphs[i].text
        if text.strip():
            body_texts.append(text)

    # --- Component 1: Three Kim bibliography entries present (0.30 points) ---
    try:
        kim_refs = [r for r in ref_paragraphs if r.startswith('Kim, J.') or r.startswith('Kim, J,')]
        kim_count = len(kim_refs)
        if kim_count >= 3:
            print(f"PASS: Component 1 -- Found {kim_count} Kim bibliography entries (0.30 pts)")
            total_score += 0.30
        elif kim_count > 0:
            partial = round(0.10 * kim_count, 2)
            print(f"PARTIAL: Component 1 -- Found {kim_count}/3 Kim entries ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- No Kim bibliography entries found")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # --- Component 2: Correct titles and years in Kim entries (0.20 points) ---
    try:
        found_titles = {
            'social_media': False,
            'digital_wellbeing': False,
            'screen_time': False,
        }
        for ref in kim_refs:
            ref_lower = ref.lower()
            if 'social media' in ref_lower and 'adolescent mental health' in ref_lower:
                found_titles['social_media'] = True
            if 'digital wellbeing' in ref_lower or 'digital well-being' in ref_lower:
                if 'intervention' in ref_lower:
                    found_titles['digital_wellbeing'] = True
            if ('long-term' in ref_lower or 'long term' in ref_lower) and 'screen time' in ref_lower:
                found_titles['screen_time'] = True

        titles_found = sum(found_titles.values())
        if titles_found == 3:
            print(f"PASS: Component 2 -- All 3 Kim entries have correct titles (0.20 pts)")
            total_score += 0.20
        elif titles_found > 0:
            partial = round(0.20 * titles_found / 3, 2)
            print(f"PARTIAL: Component 2 -- {titles_found}/3 correct titles ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- No Kim entries match expected titles")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # --- Component 3: 2019 entries disambiguated with 'a' and 'b' in bibliography (0.20 points) ---
    try:
        has_2019a = False
        has_2019b = False
        for ref in kim_refs:
            if re.search(r'\(?\s*2019a\s*\)?\.?', ref):
                has_2019a = True
            if re.search(r'\(?\s*2019b\s*\)?\.?', ref):
                has_2019b = True

        if has_2019a and has_2019b:
            print(f"PASS: Component 3 -- 2019a and 2019b disambiguated in bibliography (0.20 pts)")
            total_score += 0.20
        elif has_2019a or has_2019b:
            print(f"PARTIAL: Component 3 -- Only one of 2019a/2019b found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 -- No 2019a/2019b disambiguation in bibliography")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # --- Component 4: In-text citations in document body (0.20 points) ---
    try:
        body_full = ' '.join(body_texts)
        has_cite_2019a = bool(re.search(r'Kim\s*[\(,]\s*2019a\s*[\),]', body_full))
        has_cite_2019b = bool(re.search(r'Kim\s*[\(,]\s*2019b\s*[\),]', body_full))
        has_cite_2020 = bool(re.search(r'Kim\s*[\(,]\s*2020\s*[\),]', body_full))

        cite_count = sum([has_cite_2019a, has_cite_2019b, has_cite_2020])
        if cite_count == 3:
            print(f"PASS: Component 4 -- All 3 in-text citations found (0.20 pts)")
            total_score += 0.20
        elif cite_count > 0:
            partial = round(0.20 * cite_count / 3, 2)
            print(f"PARTIAL: Component 4 -- {cite_count}/3 in-text citations ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 -- No Kim in-text citations found in body")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # --- Component 5: In-text citation suffixes match bibliography (0.10 points) ---
    try:
        # Verify that: 2019a cites social media paper, 2019b cites digital wellbeing paper
        # Check that the in-text citation 2019a appears near "social media" context
        # and 2019b appears near "digital wellbeing" or "intervention" context
        # And that bibliography 2019a = social media, 2019b = digital wellbeing
        bib_2019a_social = False
        bib_2019b_wellbeing = False
        for ref in kim_refs:
            ref_lower = ref.lower()
            if '2019a' in ref_lower and 'social media' in ref_lower:
                bib_2019a_social = True
            if '2019b' in ref_lower and ('digital wellbeing' in ref_lower or 'digital well-being' in ref_lower):
                bib_2019b_wellbeing = True

        if bib_2019a_social and bib_2019b_wellbeing and has_cite_2019a and has_cite_2019b:
            print(f"PASS: Component 5 -- Citation suffixes are consistent (0.10 pts)")
            total_score += 0.10
        else:
            details = []
            if not bib_2019a_social:
                details.append("bib 2019a not matched to social media paper")
            if not bib_2019b_wellbeing:
                details.append("bib 2019b not matched to digital wellbeing paper")
            if not has_cite_2019a:
                details.append("no in-text 2019a citation")
            if not has_cite_2019b:
                details.append("no in-text 2019b citation")
            print(f"FAIL: Component 5 -- {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
