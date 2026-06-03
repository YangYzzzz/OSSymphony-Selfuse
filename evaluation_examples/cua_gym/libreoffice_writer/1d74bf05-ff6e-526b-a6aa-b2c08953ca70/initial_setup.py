"""
Initial Setup: Psychology paper with bibliography (no Kim entries)
Task ID: writer_bs_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('The Impact of Digital Technology on Adolescent Psychological Development', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author info ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Dr. Emily R. Thompson')
    run.font.size = Pt(12)
    run2 = author_para.add_run('\nDepartment of Psychology, University of Western Ontario')
    run2.font.size = Pt(10)
    run2.font.italic = True

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph(
        'This paper examines the multifaceted relationship between digital technology use '
        'and psychological development in adolescents aged 12-18. Drawing on recent longitudinal '
        'studies and cross-sectional analyses, we explore how varying patterns of social media '
        'engagement, screen time, and digital communication affect emotional regulation, '
        'self-esteem, and interpersonal relationships. Our findings suggest that the relationship '
        'between technology use and mental health outcomes is more nuanced than popular discourse '
        'suggests, with both protective and risk factors emerging from digital engagement.'
    )
    abstract.paragraph_format.space_after = Pt(12)

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The rapid proliferation of digital technologies has fundamentally altered the social '
        'landscape in which contemporary adolescents develop. Smartphones, social media platforms, '
        'and instant messaging applications have become integral to the daily lives of young people '
        'across the globe (Anderson & Jiang, 2018). By 2023, an estimated 95% of teenagers in '
        'developed nations had access to a smartphone, and approximately 72% reported using at '
        'least one social media platform daily (Pew Research Center, 2023).'
    )
    doc.add_paragraph(
        'The implications of this digital immersion for adolescent psychological development have '
        'become a subject of intense scholarly debate. Early research tended to adopt a predominantly '
        'negative framing, emphasizing associations between screen time and outcomes such as '
        'depression, anxiety, and reduced academic performance (Twenge & Campbell, 2018). However, '
        'more recent investigations have painted a more complex picture, identifying potential benefits '
        'alongside risks (Odgers & Jensen, 2020).'
    )
    doc.add_paragraph(
        'This paper aims to provide a comprehensive review of the current state of knowledge '
        'regarding digital technology and adolescent mental health. We examine evidence from '
        'multiple methodological approaches and consider the moderating factors that may determine '
        'whether technology use is beneficial or harmful for developing minds.'
    )

    # --- Literature Review ---
    doc.add_heading('2. Literature Review', level=1)

    doc.add_heading('2.1 Social Media and Emotional Wellbeing', level=2)
    doc.add_paragraph(
        'Research on social media and adolescent emotional wellbeing has yielded mixed results. '
        'A meta-analysis by Huang (2017) found small but significant negative correlations between '
        'social media use and subjective wellbeing. Conversely, Valkenburg and Peter (2011) '
        'demonstrated that online social interactions could enhance friendship quality and reduce '
        'loneliness for socially anxious adolescents.'
    )
    doc.add_paragraph(
        'The type of social media engagement appears to matter considerably. Passive consumption '
        'of content has been consistently associated with downward social comparison and decreased '
        'life satisfaction (Verduyn et al., 2015). Active engagement, including posting, commenting, '
        'and direct messaging, shows a more positive association with psychological outcomes '
        '(Burke & Kraut, 2016).'
    )

    doc.add_heading('2.2 Screen Time and Cognitive Development', level=2)
    doc.add_paragraph(
        'The relationship between total screen time and cognitive development has been examined '
        'in several large-scale studies. Walsh et al. (2018) analyzed data from over 4,500 '
        'children and found that exceeding two hours of recreational screen time per day was '
        'associated with lower scores on cognitive assessments. However, Przybylski and Weinstein '
        '(2017) proposed a "Goldilocks hypothesis," suggesting that moderate technology use may '
        'actually be beneficial for adolescent wellbeing.'
    )

    doc.add_heading('2.3 Digital Communication and Interpersonal Skills', level=2)
    doc.add_paragraph(
        'The impact of digital communication on interpersonal skill development remains contested. '
        'Uhls et al. (2014) found that preteens who spent five days without access to screens showed '
        'significant improvements in reading nonverbal emotional cues compared to a control group. '
        'Meanwhile, Davis (2012) argued that digital communication platforms provide valuable '
        'practice spaces for identity exploration and social skill development.'
    )

    # --- Methodology ---
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph(
        'This review synthesizes findings from 47 peer-reviewed studies published between 2015 and '
        '2024. We employed a systematic search strategy using PsycINFO, PubMed, and Web of Science '
        'databases. Inclusion criteria required empirical studies with adolescent samples (ages 12-18), '
        'quantitative measures of digital technology use, and validated psychological outcome measures. '
        'Studies were assessed for methodological quality using the Newcastle-Ottawa Scale for '
        'observational studies and the Cochrane Risk of Bias tool for randomized controlled trials.'
    )

    # --- Discussion ---
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph(
        'Our analysis reveals several key themes that warrant discussion. First, the relationship '
        'between digital technology and adolescent mental health is not uniformly negative. The evidence '
        'suggests that context, content, and individual differences play crucial moderating roles. '
        'Second, methodological limitations in existing research, including reliance on cross-sectional '
        'designs and self-report measures, constrain the strength of causal inferences that can be drawn.'
    )
    doc.add_paragraph(
        'A particularly important finding is the distinction between active and passive technology '
        'use. While passive scrolling through curated social media feeds appears to be associated '
        'with negative outcomes, active engagement in meaningful digital interactions may support '
        'social development. This distinction has important implications for both parents and '
        'policymakers seeking evidence-based guidance on adolescent technology use.'
    )

    # --- Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        'The relationship between digital technology and adolescent psychological development '
        'is complex and multifaceted. Rather than adopting a simplistic stance of technology as '
        'either universally harmful or benign, researchers and practitioners must attend to the '
        'specific patterns of use, the developmental stage of the adolescent, and the broader '
        'social context in which technology is embedded. Future research should prioritize '
        'longitudinal designs, objective measures of technology use, and investigation of '
        'individual difference moderators.'
    )

    # --- References ---
    doc.add_heading('References', level=1)

    references = [
        'Anderson, M., & Jiang, J. (2018). Teens, social media & technology 2018. Pew Research Center.',
        'Burke, M., & Kraut, R. (2016). The relationship between Facebook use and well-being depends on communication type and tie strength. Journal of Computer-Mediated Communication, 21(4), 265-281.',
        'Davis, K. (2012). Friendship 2.0: Adolescents\' experiences of belonging and self-disclosure online. Journal of Adolescence, 35(6), 1527-1536.',
        'Huang, C. (2017). Time spent on social network sites and psychological well-being: A meta-analysis. Cyberpsychology, Behavior, and Social Networking, 20(6), 346-354.',
        'Odgers, C. L., & Jensen, M. R. (2020). Annual research review: Adolescent mental health in the digital age. Journal of Child Psychology and Psychiatry, 61(3), 336-348.',
        'Pew Research Center. (2023). Teens, social media and technology 2023. Pew Research Center.',
        'Przybylski, A. K., & Weinstein, N. (2017). A large-scale test of the Goldilocks hypothesis. Psychological Science, 28(2), 204-215.',
        'Twenge, J. M., & Campbell, W. K. (2018). Associations between screen time and lower psychological well-being among children and adolescents. Preventive Medicine Reports, 12, 271-283.',
        'Uhls, Y. T., Michikyan, M., Morris, J., Garcia, D., Small, G. W., Zgourou, E., & Greenfield, P. M. (2014). Five days at outdoor education camp without screens improves preteen skills with nonverbal emotion cues. Computers in Human Behavior, 39, 387-392.',
        'Valkenburg, P. M., & Peter, J. (2011). Online communication among adolescents: An integrated model of its attraction, opportunities, and risks. Journal of Adolescent Health, 48(2), 121-127.',
        'Verduyn, P., Lee, D. S., Park, J., Shablack, H., Orvell, A., Bayer, J., ... & Kross, E. (2015). Passive Facebook usage undermines affective well-being. Journal of Experimental Psychology: General, 144(2), 480-488.',
        'Walsh, J. J., Barnes, J. D., Cameron, J. D., Goldfield, G. S., Chaput, J. P., Gunnell, K. E., ... & Tremblay, M. S. (2018). Associations between 24 hour movement behaviours and global cognition in US children. The Lancet Child & Adolescent Health, 2(11), 783-791.',
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
