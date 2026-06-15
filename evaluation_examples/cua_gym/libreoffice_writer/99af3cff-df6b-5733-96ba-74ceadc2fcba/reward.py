"""
Reward Script: Add page numbers centered in footer with 'Page X of Y' format
Task ID: writer_rd_007
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Footer is enabled (not linked to previous)
  Component 2 (0.35): Footer contains 'Page X of Y' field structure (PAGE + NUMPAGES fields)
  Component 3 (0.20): Footer paragraph is center-aligned
  Component 4 (0.25): Footer font is Liberation Serif 10pt
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_007'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one section
    if len(doc.sections) == 0:
        print("FAIL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    footer = section.footer

    # Component 1: Footer is enabled — not linked to previous (0.20 points)
    # Initial state has footer linked_to_previous=True (default/empty).
    # Golden state has it set to False (footer is explicitly defined).
    try:
        is_linked = footer.is_linked_to_previous
        if is_linked is False:
            print(f"PASS: Component 1 — Footer is enabled (is_linked_to_previous=False) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Footer is_linked_to_previous={is_linked}, expected False")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Footer contains 'Page X of Y' structure with PAGE and NUMPAGES fields (0.35 points)
    # Initial state: no footer text, no field codes.
    # Golden state: footer text is "Page  of " with PAGE and NUMPAGES instrText fields.
    try:
        footer_paras = footer.paragraphs
        if len(footer_paras) == 0:
            print("FAIL: Component 2 — No paragraphs in footer")
        else:
            fp = footer_paras[0]
            fp_xml = fp._element.xml

            # Check for PAGE field code
            instr_texts = [instr.text.strip() for instr in fp._element.findall('.//' + qn('w:instrText'))]
            has_page_field = any('PAGE' in t and 'NUMPAGES' not in t for t in instr_texts)
            has_numpages_field = any('NUMPAGES' in t for t in instr_texts)

            # Check for "Page" and "of" text in runs
            run_texts = [r.text for r in fp.runs if r.text]
            combined_text = ''.join(run_texts).strip()
            # Use the full (untrimmed) combined text for pattern matching
            raw_text = ''.join(run_texts)
            has_page_text = 'Page' in raw_text or 'page' in raw_text
            has_of_text = ' of ' in raw_text or ' of' in raw_text

            sub_score = 0.0
            details = []
            if has_page_field:
                sub_score += 0.10
                details.append("PAGE field found")
            else:
                details.append("PAGE field MISSING")
            if has_numpages_field:
                sub_score += 0.10
                details.append("NUMPAGES field found")
            else:
                details.append("NUMPAGES field MISSING")
            if has_page_text and has_of_text:
                sub_score += 0.15
                details.append(f"Text structure OK: '{combined_text}'")
            else:
                details.append(f"Text structure MISSING: got '{combined_text}', need 'Page ... of ...'")

            if sub_score >= 0.35:
                print(f"PASS: Component 2 — 'Page X of Y' structure verified ({'; '.join(details)}) (0.35 pts)")
                total_score += 0.35
            elif sub_score > 0:
                print(f"PARTIAL: Component 2 — {'; '.join(details)} ({sub_score} pts)")
                total_score += sub_score
            else:
                print(f"FAIL: Component 2 — {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Footer paragraph is center-aligned (0.20 points)
    # Initial state: alignment is None (default/left).
    # Golden state: alignment is CENTER.
    try:
        footer_paras = footer.paragraphs
        if len(footer_paras) == 0:
            print("FAIL: Component 3 — No paragraphs in footer")
        else:
            fp = footer_paras[0]
            alignment = fp.paragraph_format.alignment
            if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                print(f"PASS: Component 3 — Footer is center-aligned (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Footer alignment is {alignment}, expected CENTER")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Footer font is Liberation Serif 10pt (0.25 points)
    # Initial state: no runs in footer.
    # Golden state: runs with Liberation Serif, size=127000 EMU (10pt).
    try:
        footer_paras = footer.paragraphs
        if len(footer_paras) == 0:
            print("FAIL: Component 4 — No paragraphs in footer")
        else:
            fp = footer_paras[0]
            runs_with_font = [r for r in fp.runs if r.font.name is not None or r.font.size is not None]
            if len(runs_with_font) == 0:
                print("FAIL: Component 4 — No runs with font properties in footer")
            else:
                correct_name = 0
                correct_size = 0
                total_checked = len(runs_with_font)
                for r in runs_with_font:
                    if r.font.name and 'Liberation Serif' in r.font.name:
                        correct_name += 1
                    if r.font.size is not None and abs(r.font.size.pt - 10.0) < 0.5:
                        correct_size += 1

                name_ok = correct_name >= total_checked * 0.5
                size_ok = correct_size >= total_checked * 0.5

                if name_ok and size_ok:
                    print(f"PASS: Component 4 — Font is Liberation Serif 10pt ({correct_name}/{total_checked} name, {correct_size}/{total_checked} size) (0.25 pts)")
                    total_score += 0.25
                elif name_ok or size_ok:
                    partial = 0.125
                    print(f"PARTIAL: Component 4 — name_ok={name_ok}, size_ok={size_ok} ({partial} pts)")
                    total_score += partial
                else:
                    sample_names = set(r.font.name for r in runs_with_font if r.font.name)
                    sample_sizes = set(r.font.size.pt for r in runs_with_font if r.font.size)
                    print(f"FAIL: Component 4 — Font names: {sample_names}, sizes: {sample_sizes}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
