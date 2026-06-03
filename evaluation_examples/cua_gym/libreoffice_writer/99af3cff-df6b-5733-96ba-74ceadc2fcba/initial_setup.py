"""
Initial Setup: Add page numbers centered in footer (Page X of Y)
Task ID: writer_rd_007
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default page style margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Ensure no footer is enabled (default Document has no footer content)
    # Do NOT add any header or footer

    # --- Page 1: Title page ---
    title = doc.add_heading('Quarterly Business Review', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(16)
    run.font.name = 'Liberation Serif'

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Q4 2025 - Prepared by Strategic Planning Division')
    run.font.size = Pt(12)
    run.font.name = 'Liberation Serif'

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Confidential - Internal Use Only').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Page 2: Executive Summary ---
    doc.add_page_break()
    doc.add_heading('Executive Summary', level=1)

    doc.add_paragraph(
        'Meridian Technologies concluded Q4 2025 with strong performance across all major '
        'business segments. Total revenue reached $147.3 million, representing a 12.4% '
        'year-over-year increase. The company successfully launched three new product lines '
        'and expanded operations into two additional markets in Southeast Asia.'
    )

    doc.add_paragraph(
        'Operating expenses were well-managed at $98.6 million, resulting in an operating '
        'margin of 33.1%. The engineering division completed 94% of planned deliverables, '
        'while customer satisfaction scores improved to 4.7 out of 5.0, up from 4.3 in Q3.'
    )

    doc.add_paragraph(
        'Key highlights for the quarter include the successful deployment of the Aurora '
        'platform, which onboarded 2,300 enterprise clients in its first 60 days. The '
        'partnership with NovaTech Solutions contributed $18.2 million in joint revenue, '
        'exceeding the projected $14.5 million target by 25.5%.'
    )

    doc.add_paragraph(
        'Looking ahead to Q1 2026, the leadership team has approved a $23 million investment '
        'in artificial intelligence capabilities, targeting the healthcare and financial '
        'services verticals. The hiring plan calls for 145 new positions across engineering, '
        'sales, and customer success departments.'
    )

    # --- Page 3: Financial Overview ---
    doc.add_page_break()
    doc.add_heading('Financial Overview', level=1)

    doc.add_heading('Revenue Breakdown', level=2)
    doc.add_paragraph(
        'The following table summarizes revenue by business segment for Q4 2025:'
    )

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['Segment', 'Q4 Revenue ($M)', 'Q3 Revenue ($M)', 'Growth (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Enterprise Software', '58.2', '52.1', '11.7'],
        ['Cloud Services', '41.7', '35.9', '16.2'],
        ['Professional Services', '22.4', '21.0', '6.7'],
        ['Hardware Solutions', '15.8', '14.3', '10.5'],
        ['Licensing & Support', '9.2', '8.4', '9.5'],
        ['Total', '147.3', '131.7', '11.8'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'Enterprise Software remains our largest revenue driver, accounting for 39.5% of '
        'total revenue. Cloud Services showed the strongest growth at 16.2%, driven by '
        'increased adoption of our managed infrastructure offering among mid-market clients.'
    )

    # --- Page 4: Product Development ---
    doc.add_page_break()
    doc.add_heading('Product Development', level=1)

    doc.add_heading('Aurora Platform Launch', level=2)
    doc.add_paragraph(
        'The Aurora Platform, our next-generation enterprise collaboration suite, launched on '
        'October 15, 2025. Key metrics from the first 75 days include:'
    )

    bullets = [
        '2,300 enterprise clients onboarded (target: 1,800)',
        '98.7% uptime across all regions',
        'Average response time of 145ms (SLA requirement: 200ms)',
        'Net Promoter Score of 72 among early adopters',
        '340 feature requests catalogued and prioritized',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('Quantum Analytics Module', level=2)
    doc.add_paragraph(
        'The Quantum Analytics module entered beta testing with 150 selected clients on '
        'November 1, 2025. Initial feedback has been overwhelmingly positive, with 89% of '
        'beta testers indicating they would upgrade to the premium tier upon general '
        'availability. The module processes an average of 4.2 billion data points per day '
        'across all beta instances.'
    )

    doc.add_paragraph(
        'Engineering resolved 247 reported issues during the beta period, with a mean time '
        'to resolution of 3.2 business days. Critical and high-severity issues were addressed '
        'within 24 hours per our service level agreement.'
    )

    # --- Page 5: Market Expansion ---
    doc.add_page_break()
    doc.add_heading('Market Expansion', level=1)

    doc.add_heading('Southeast Asia Operations', level=2)
    doc.add_paragraph(
        'Meridian Technologies established regional offices in Jakarta, Indonesia and '
        'Ho Chi Minh City, Vietnam during Q4 2025. The Southeast Asian market represents a '
        '$4.8 billion addressable opportunity in enterprise software, growing at 18% annually.'
    )

    doc.add_paragraph(
        'Our initial go-to-market strategy focuses on financial services and manufacturing '
        'verticals, where we have strong reference accounts from our APAC operations. '
        'The regional team of 35 professionals includes sales engineers, customer success '
        'managers, and technical consultants fluent in Bahasa Indonesia, Vietnamese, and English.'
    )

    doc.add_heading('European Market Performance', level=2)
    doc.add_paragraph(
        'European operations contributed $32.4 million in Q4 revenue, a 14.1% increase over '
        'Q3. The DACH region (Germany, Austria, Switzerland) accounted for 45% of European '
        'revenue, followed by the UK at 28% and France at 15%. We onboarded 18 new enterprise '
        'clients in Europe during the quarter, including three Fortune Global 500 companies.'
    )

    doc.add_paragraph(
        'Compliance with the EU Digital Services Act and updated GDPR requirements was '
        'achieved ahead of schedule, positioning Meridian favorably against competitors who '
        'are still working toward full compliance.'
    )

    # --- Page 6: Human Resources ---
    doc.add_page_break()
    doc.add_heading('Human Resources & Organizational Development', level=1)

    doc.add_paragraph(
        'Meridian Technologies ended Q4 2025 with 2,847 full-time employees across 14 '
        'global offices. The company hired 198 new team members during the quarter, with '
        'a 90-day retention rate of 96.3%.'
    )

    doc.add_heading('Employee Satisfaction', level=2)
    doc.add_paragraph(
        'The annual employee engagement survey, conducted in November 2025, showed an '
        'overall satisfaction score of 4.2 out of 5.0. Key findings include:'
    )

    bullets2 = [
        'Work-life balance: 4.4/5.0 (up from 3.9 in 2024)',
        'Career development opportunities: 4.1/5.0',
        'Compensation and benefits: 3.8/5.0',
        'Management effectiveness: 4.3/5.0',
        'Company culture and values: 4.5/5.0',
    ]
    for b in bullets2:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'The learning and development team facilitated 1,240 training hours across the '
        'organization, including specialized programs in cloud architecture, AI/ML engineering, '
        'and leadership development. The internal mobility program saw 67 employees transition '
        'to new roles within the company.'
    )

    # --- Page 7: Risk Assessment ---
    doc.add_page_break()
    doc.add_heading('Risk Assessment & Mitigation', level=1)

    doc.add_paragraph(
        'The enterprise risk management committee identified and monitored 23 risk items '
        'during Q4 2025. The following represent the highest-priority risks and their '
        'current mitigation status:'
    )

    risk_table = doc.add_table(rows=6, cols=3)
    risk_table.style = 'Table Grid'
    risk_headers = ['Risk Category', 'Severity', 'Mitigation Status']
    for i, h in enumerate(risk_headers):
        cell = risk_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    risks = [
        ['Cybersecurity threats', 'High', 'Active monitoring; SOC team expanded'],
        ['Supply chain disruption', 'Medium', 'Dual-sourcing strategy implemented'],
        ['Regulatory compliance (APAC)', 'Medium', 'Legal review in progress'],
        ['Key personnel retention', 'Low', 'Retention packages for top 50 leaders'],
        ['Currency fluctuation', 'Medium', 'Hedging strategy for EUR/USD/SGD'],
    ]
    for r, row_data in enumerate(risks, 1):
        for c, val in enumerate(row_data):
            risk_table.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'The cybersecurity team conducted two red team exercises during the quarter, '
        'identifying and remediating 14 vulnerabilities before they could be exploited. '
        'The company maintained its SOC 2 Type II certification and completed the ISO 27001 '
        'recertification audit with zero non-conformities.'
    )

    # --- Page 8: Strategic Outlook ---
    doc.add_page_break()
    doc.add_heading('Strategic Outlook for 2026', level=1)

    doc.add_paragraph(
        'Meridian Technologies enters 2026 with strong momentum and a clear strategic vision. '
        'The board of directors has approved the following strategic priorities for the fiscal year:'
    )

    priorities = [
        'Achieve $620 million in annual revenue (18% growth target)',
        'Launch the Meridian AI Suite for healthcare and financial services',
        'Expand to 5 additional markets in Asia-Pacific and Latin America',
        'Increase recurring revenue to 75% of total revenue mix',
        'Obtain FedRAMP High authorization for government sector entry',
        'Grow workforce to 3,500 employees by year-end 2026',
    ]
    for i, p in enumerate(priorities, 1):
        doc.add_paragraph(f'{p}', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph(
        'The capital expenditure budget for 2026 is set at $45 million, with primary '
        'allocations toward data center expansion ($18M), product development ($15M), and '
        'market entry costs ($12M). The company expects to maintain a free cash flow margin '
        'of at least 20% throughout the year.'
    )

    doc.add_paragraph(
        'In conclusion, Q4 2025 demonstrated Meridian Technologies\' ability to execute on '
        'growth initiatives while maintaining operational discipline. The management team is '
        'confident in the company\'s ability to achieve its 2026 objectives and deliver '
        'sustained value to shareholders, employees, and clients.'
    )

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
