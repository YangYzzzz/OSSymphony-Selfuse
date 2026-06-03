"""
Initial Setup: Novel manuscript with single-sided layout and simple headers
Task ID: writer_rd_082
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_082'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# Chapter titles for a realistic novel
CHAPTER_TITLES = [
    "The Awakening",
    "Shadows of the Past",
    "A Stranger in Town",
    "The Hidden Letter",
    "Crossroads",
    "Into the Storm",
    "Revelations",
    "The Betrayal",
    "A Glimmer of Hope",
    "The Final Confrontation",
    "Aftermath",
    "New Beginnings",
]

# Realistic novel paragraphs for each chapter (varied content)
CHAPTER_CONTENT = [
    [
        "The morning sun crept through the curtains of the old Victorian house on Maple Street, casting long golden fingers across the hardwood floor. Eleanor Blackwood stirred in her chair, the worn leather creaking beneath her as she shifted. She had fallen asleep reading again, the heavy volume of maritime history still open across her lap, its pages yellowed with age.",
        "Outside, the town of Havenport was already coming alive. The fishermen had been up since before dawn, their boats cutting dark silhouettes against the brightening horizon. Mrs. Patterson from the bakery was arranging her display of fresh sourdough and cinnamon rolls, the sweet aroma drifting down the cobblestone lane like an invitation.",
        "Eleanor closed the book carefully and placed it on the side table beside her cold cup of tea. Today was different from the quiet days that had preceded it. Today marked exactly thirty years since the disappearance of the Meridian, her father's research vessel, and with it, Captain Thomas Blackwood himself.",
        "She walked to the window and looked out at the harbor. The water was calm, almost glass-like, reflecting the pastel sky. Somewhere beneath those deceptively peaceful waves lay the answers she had spent half her life searching for.",
    ],
    [
        "The Blackwood family had once been the most prominent in Havenport. Thomas Blackwood, Eleanor's father, had been a renowned marine archaeologist whose discoveries had put their small coastal town on the map. His expeditions had uncovered shipwrecks dating back to the colonial era, each one yielding artifacts that filled the town's modest museum.",
        "But that all changed on the night of October 14th, 1995. The Meridian had set out on what Thomas called his most important expedition yet. He had been secretive about the details, sharing only with his first mate, Jorge Castillo, and his research partner, Dr. Helen Weaver. The three of them, along with a crew of twelve, sailed into a gathering storm and were never seen again.",
        "The Coast Guard search lasted three weeks. They found debris scattered across a fifty-mile radius, but no bodies, no black box, and no explanation for why an experienced captain would sail directly into a Category 3 hurricane. The official report concluded mechanical failure compounded by severe weather conditions.",
        "Eleanor had been twelve years old. She remembered standing at the dock with her mother, watching the search boats return empty-handed day after day. Her mother, Catherine, had aged a decade in those three weeks. Within a year, she had sold the family estate and moved them to a small apartment above the hardware store on Main Street.",
    ],
    [
        "The stranger arrived on the morning ferry from Portland, carrying nothing but a worn leather satchel and an air of quiet determination. He was tall, with weathered skin that spoke of years spent outdoors and eyes the color of sea glass. He wore a canvas jacket despite the warmth of the late September morning.",
        "Margaret Chen, who ran the Harbor View Inn, was the first to meet him. She would later tell anyone who would listen that he had asked specifically about the Blackwood family, which was unusual enough to set the town's gossip network buzzing within the hour.",
        "His name, he said, was Daniel Reeves. He claimed to be a journalist working on a feature about unsolved maritime mysteries for National Geographic. It was a plausible enough cover story, and Margaret, who prided herself on being a gracious hostess, showed him to the best room overlooking the harbor.",
        "But Daniel Reeves was not a journalist. The press credentials in his wallet were forgeries, expertly made but forgeries nonetheless. The real purpose of his visit to Havenport was far more personal and far more dangerous than anyone in the sleepy coastal town could have imagined.",
    ],
    [
        "Eleanor found the letter on a Tuesday afternoon, tucked between the pages of her father's personal journal. She had been going through his papers again, as she did every year around the anniversary, looking for anything she might have missed in her previous searches.",
        "The envelope was unmarked, sealed with red wax bearing an unfamiliar crest: a compass rose overlaid with what appeared to be a serpent. Inside was a single sheet of heavy cream paper, covered in her father's precise handwriting. The date at the top read September 28, 1995, just sixteen days before the Meridian's final voyage.",
        "My dearest Eleanor, the letter began. If you are reading this, then I did not return as I promised. I am sorry. What I am about to do is dangerous, but it is necessary. The Cartwright Collection is real, and I believe I have found it. The coordinates are encoded in the margin notes of Hargrove's Atlas, third edition, which you will find in my study safe. The combination is your mother's birthday, reversed.",
        "Eleanor's hands trembled as she read. The Cartwright Collection was legendary among maritime historians: a cache of artifacts from the lost colony of Roanoke, supposedly hidden by Sir Francis Drake aboard a vessel that sank somewhere off the New England coast in 1587. Most scholars dismissed it as myth. Her father, apparently, had not.",
    ],
    [
        "The discovery of the letter forced Eleanor to a crossroads she had been avoiding for decades. She could continue her quiet life as Havenport's librarian, cataloging books and hosting the weekly children's reading hour, pretending that the mysteries of the past held no power over her present.",
        "Or she could open her father's safe, decode the coordinates, and finish what he had started. The choice should have been simple for a woman who had spent thirty years seeking answers. But Eleanor understood something now that her younger self had not: pursuing the truth might reveal things she was not prepared to face.",
        "She sat in the library after closing time, the letter spread on the oak desk before her, and weighed her options. The building was silent except for the ticking of the grandfather clock in the reading room and the distant sound of waves breaking against the seawall.",
        "Her phone buzzed. A text from her friend and fellow researcher, Professor James Whitfield at the University of Maine: Saw the strangest thing today. Someone accessed the restricted maritime archives asking about the Cartwright legend. Thought you should know. Eleanor felt the hairs on the back of her neck rise.",
    ],
    [
        "Hurricane Nadine formed unexpectedly in the mid-Atlantic on the first day of October, defying the meteorological models that had predicted a quiet end to the season. Within forty-eight hours, it had strengthened to a Category 2 and was tracking directly toward the New England coast.",
        "For Havenport, hurricanes were nothing new. The town had weathered dozens of storms over its three-hundred-year history, and the residents knew the drill: board the windows, stock up on supplies, and wait it out. But this storm felt different. The barometric pressure was dropping faster than any of the old-timers could remember.",
        "Eleanor watched the Weather Channel's tracking graphic with growing unease. The projected path of Nadine followed almost exactly the route her father had taken thirty years ago. Coincidence, probably. But the timing, just as she had discovered his letter, felt less like chance and more like a sign.",
        "Daniel Reeves appeared at the library that afternoon, his journalist cover story wearing thin. He asked Eleanor directly about her father's last expedition, and for the first time, she saw something in his eyes that went beyond professional curiosity. It was fear.",
    ],
    [
        "The truth, when it finally emerged, was more extraordinary than even Eleanor had imagined. Daniel Reeves was not a journalist. He was the son of Jorge Castillo, her father's first mate. And he had in his possession the other half of the puzzle.",
        "Jorge Castillo had survived the wreck of the Meridian. Badly injured and suffering from hypothermia, he had been picked up by a Portuguese fishing vessel three days after the storm. But instead of returning to Havenport, he had disappeared, assuming a new identity and spending the next three decades in hiding.",
        "The reason, Daniel explained over cups of coffee in Eleanor's cramped apartment, was that the sinking of the Meridian had not been an accident. Someone had sabotaged the vessel, and Jorge believed that the same person or organization would kill anyone who knew the truth about the Cartwright Collection.",
        "Jorge had died two years ago, cancer, but not before passing his knowledge to his son. Daniel had spent those two years verifying his father's claims, and everything checked out. The Cartwright Collection was not just real; it was worth an estimated four hundred million dollars.",
    ],
    [
        "Dr. Helen Weaver lived in a renovated lighthouse on Cape Ashford, twenty miles north of Havenport. She had retired from academia a decade ago, ostensibly to write her memoirs, though the book remained perpetually unfinished. Her neighbors knew her as a quiet, private woman who kept cats and grew spectacular roses.",
        "Eleanor and Daniel drove up the coast road on a gray morning, the ocean churning to their left. Eleanor had not seen Helen since the memorial service for the Meridian's crew. Helen had spoken briefly, her voice steady but her eyes hollow, and then had disappeared from Havenport entirely.",
        "The lighthouse door opened before they could knock. Helen stood there, smaller than Eleanor remembered, her silver hair pulled back in a severe bun. She looked at Eleanor for a long moment, then at Daniel, and her face went white.",
        "You look just like your father, she whispered to Daniel. Then she stepped aside and said, Come in. I have been expecting someone to come asking questions for thirty years. It is almost a relief. The betrayal Eleanor was about to uncover would shake her understanding of everything she thought she knew about her father's legacy.",
    ],
    [
        "In the days following Helen's confession, Eleanor found herself unable to sleep. The revelation that Helen had been coerced into providing the expedition's coordinates to a private collector named Aldric Voss had been shocking enough. But learning that her father had suspected the betrayal and sailed anyway was almost unbearable.",
        "Thomas Blackwood had known the risks. His letter to Eleanor proved that. He had prepared for the possibility that he would not return, and he had taken steps to ensure that the location of the Cartwright Collection would survive even if he did not. What he could not have anticipated was that the secret would remain buried for three decades.",
        "But there was a glimmer of hope in the darkness. If Jorge Castillo had survived, perhaps others had too. Daniel's research suggested that at least three other crew members had been rescued by the same Portuguese vessel, though their trails went cold in Lisbon. A detective agency in Portugal was following the leads.",
        "Eleanor sat on the beach at sunset, watching the waves roll in, each one erasing the footprints of the last visitor. The ocean kept its secrets well, but not forever. Tides changed, storms shifted the seabed, and what was lost could always be found by someone patient enough to keep looking.",
    ],
    [
        "The confrontation with Aldric Voss took place not in some dramatic seaside location, but in the sterile conference room of his Manhattan law firm on the forty-seventh floor of a glass tower. Voss was eighty-three years old, wheelchair-bound, and attached to a portable oxygen tank. He did not look like a man capable of destroying lives.",
        "Eleanor had expected to feel rage when she finally faced the person responsible for her father's disappearance. Instead, she felt an overwhelming sadness. Voss was a collector, obsessed with possessing things that belonged to history, willing to kill for objects that should have been shared with the world.",
        "The meeting lasted four hours. Voss's lawyers tried to shut it down multiple times, but the old man waved them away. He wanted to talk. Perhaps, at the end of his life, he too was tired of secrets. He confirmed everything: the sabotage, the cover-up, the bribes to the Coast Guard investigator who had classified the incident as weather-related.",
        "But he also revealed something unexpected. He had never found the Cartwright Collection. The Meridian had gone down before reaching the coordinates, and his subsequent expeditions using the stolen data had come up empty. The collection was still out there, waiting.",
    ],
    [
        "The aftermath of the Voss confrontation played out over months. Federal prosecutors, armed with Eleanor's evidence and Daniel's documentation, opened an investigation that would eventually lead to indictments for conspiracy, maritime sabotage, and involuntary manslaughter.",
        "Havenport was thrown into the national spotlight. Camera crews from every major network descended on the quiet town, filming the harbor, the Blackwood apartment, the library where Eleanor worked. She gave exactly one interview, to a reporter from the Portland Press Herald, and then retreated into silence.",
        "The Portuguese detective agency found two of the surviving crew members. Marcus Webb was living under an assumed name in the Azores, running a small boat repair shop. He wept when Eleanor called him, apologizing for not coming forward sooner. The fear had been too great, he said. Voss had people everywhere.",
        "The second survivor, navigator Anna Rodriguez, had returned to the United States years ago and was living in Key West. She had written an account of the sabotage and the rescue, sealed it in a safety deposit box, and instructed her attorney to release it upon her death. Eleanor's investigation made that precaution unnecessary.",
    ],
    [
        "It took Eleanor Blackwood eighteen months to assemble the expedition. She used her father's encoded coordinates, cross-referenced with Jorge Castillo's notes and Anna Rodriguez's navigational records, to pinpoint the most likely location of the Cartwright Collection. The site was thirty nautical miles east of Havenport, in two hundred feet of water.",
        "The dive team included Daniel Reeves, Professor Whitfield, and a marine archaeologist from Woods Hole Oceanographic Institution named Dr. Sarah Okonkwo. They chartered a research vessel called the New Meridian, a deliberate choice that Eleanor knew her father would have appreciated.",
        "On a clear morning in July, with the sun blazing overhead and a gentle swell rolling beneath them, they descended to the ocean floor. The sonar had shown anomalies consistent with a shipwreck, and as the submersible's lights cut through the darkness, Eleanor saw the outline of wooden ribs and iron fittings emerging from the sediment.",
        "The Cartwright Collection was there, preserved in sealed bronze cases that had protected the contents from centuries of saltwater. When they brought the first case to the surface and opened it, they found maps, coins, personal letters from the colonists, and a carved wooden figure that would later be identified as a gift from the Croatoan people. Eleanor held the figure in her hands and cried, not for herself, but for her father, who had given everything to find this moment.",
        "The artifacts were donated to the Smithsonian, with a permanent loan exhibit at the Havenport Maritime Museum. A plaque at the entrance read: In memory of Captain Thomas Blackwood and the crew of the Meridian, who sought the truth at any cost. The museum opened on October 14th, exactly thirty-one years after the Meridian sailed into history.",
    ],
]


def create_initial():
    doc = Document()

    # -- Page Setup: single-sided, standard margins --
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # -- Header: "Novel Title" on all pages, regular text --
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run("Novel Title")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- Title page --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    title_run = title_para.add_run("THE GREAT ADVENTURE")
    title_run.font.size = Pt(28)
    title_run.font.name = "Times New Roman"
    title_run.bold = True

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_before = Pt(36)
    author_run = author_para.add_run("by Eleanor Blackwood")
    author_run.font.size = Pt(16)
    author_run.font.name = "Times New Roman"
    author_run.italic = True

    # Page break after title page
    doc.add_page_break()

    # -- Chapters --
    for i, (title, paragraphs) in enumerate(zip(CHAPTER_TITLES, CHAPTER_CONTENT)):
        # Chapter heading (Heading 1)
        heading = doc.add_heading(f"Chapter {i+1}: {title}", level=1)
        # Style the heading
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)

        # Chapter content
        for text in paragraphs:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(1.27)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        # Add extra paragraphs to fill pages (need ~50 pages total)
        # Each chapter needs about 3-4 pages worth of content
        filler_texts = [
            f"The events of this chapter left a lasting impression on all who witnessed them. In the weeks that followed, the residents of Havenport would speak of little else, their conversations returning again and again to the extraordinary revelations that had come to light.",
            f"Eleanor spent the evening reviewing her notes, cross-referencing dates and locations with the documents she had accumulated over the years. The picture was becoming clearer, but significant gaps remained. She knew that filling those gaps would require venturing beyond the safety of her research and into the uncertain waters of direct confrontation.",
            f"The weather turned cold that week, a premature autumn chill sweeping down from the north and stripping the last leaves from the maples along Harbor Road. The fishing fleet stayed in port, the boats rocking gently at their moorings, their crews gathered in the warmth of the Anchor Tavern sharing stories and speculation.",
            f"Daniel had set up a makeshift research station in the back room of Margaret Chen's inn, covering the walls with maps, photographs, and printouts of archived newspaper articles. The timeline was taking shape: a web of connections spanning three decades, linking maritime expeditions, private collectors, and academic institutions across two continents.",
            f"Professor Whitfield arrived from the university with a box of recently declassified Coast Guard records obtained through a Freedom of Information Act request. The documents revealed inconsistencies in the original investigation that had been overlooked or deliberately suppressed. Someone in a position of authority had ensured that certain questions were never asked.",
            f"That night, Eleanor dreamed of the ocean. In her dream, she was standing on the deck of a ship she knew to be the Meridian, though she had never set foot on it in life. The sky was dark with storm clouds, and the waves rose around the vessel like mountains. But instead of fear, she felt a profound calm, as though the sea itself was welcoming her home.",
        ]

        for text in filler_texts:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(1.27)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        # Page break between chapters (except after last)
        if i < len(CHAPTER_TITLES) - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
