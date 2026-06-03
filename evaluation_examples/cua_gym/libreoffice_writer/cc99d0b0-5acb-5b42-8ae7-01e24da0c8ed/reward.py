"""
Reward Script: Facing pages layout with differentiated even/odd headers in small caps
Task ID: writer_rd_082
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): mirrorMargins enabled
  Component 2 (0.15): evenAndOddHeaders enabled
  Component 3 (0.20): Even header = "The Great Adventure" in small caps, centered
  Component 4 (0.20): Odd header has STYLEREF field for chapter title in small caps, centered
  Component 5 (0.10): Both headers have thin bottom border
  Component 6 (0.15): Inner/outer margins = 3.0cm / 2.0cm
"""

import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_082'

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must exist and be loadable
    if not os.path.exists(file_path):
        print(f"CRITICAL: File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load document: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------------------------------------------------------------
    # Component 1: mirrorMargins enabled in settings.xml (0.20 points)
    # In initial: mirrorMargins is absent/false
    # In golden: mirrorMargins is present/true
    # ---------------------------------------------------------------
    try:
        z = zipfile.ZipFile(file_path)
        settings_xml = ET.parse(z.open('word/settings.xml'))
        root = settings_xml.getroot()
        ns = {'w': WNS}
        mirror_el = root.find('.//w:mirrorMargins', ns)
        # Element present means enabled (unless val="0" or val="false")
        mirror_enabled = False
        if mirror_el is not None:
            val = mirror_el.get(f'{{{WNS}}}val')
            # If no val attribute or val is "1"/"true"/None -> enabled
            if val is None or val in ('1', 'true'):
                mirror_enabled = True
        z.close()

        if mirror_enabled:
            print(f"PASS: Component 1 — mirrorMargins is enabled (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — mirrorMargins not enabled")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: evenAndOddHeaders enabled in settings.xml (0.15 points)
    # In initial: absent/false
    # In golden: present/true
    # ---------------------------------------------------------------
    try:
        z = zipfile.ZipFile(file_path)
        settings_xml = ET.parse(z.open('word/settings.xml'))
        root = settings_xml.getroot()
        ns = {'w': WNS}
        even_odd_el = root.find('.//w:evenAndOddHeaders', ns)
        even_odd_enabled = False
        if even_odd_el is not None:
            val = even_odd_el.get(f'{{{WNS}}}val')
            if val is None or val in ('1', 'true'):
                even_odd_enabled = True
        z.close()

        if even_odd_enabled:
            print(f"PASS: Component 2 — evenAndOddHeaders is enabled (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — evenAndOddHeaders not enabled")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Even (left) page header shows "The Great Adventure"
    #   in small caps, center-aligned (0.20 points)
    # In initial: even header is linked/empty
    # In golden: even header has text, small caps, center alignment
    # ---------------------------------------------------------------
    try:
        sec = doc.sections[0]
        even_hdr = sec.even_page_header

        # Check even header is not linked to previous (i.e., has its own content)
        even_linked = even_hdr.is_linked_to_previous

        # Get text from even header
        even_text = ''
        even_small_caps = False
        even_centered = False
        even_font_size_ok = False

        if not even_linked and even_hdr.paragraphs:
            para = even_hdr.paragraphs[0]
            even_text = para.text.strip()
            # Check alignment
            pf_align = para.paragraph_format.alignment
            if pf_align is not None and pf_align == 1:  # CENTER
                even_centered = True
            # Also check XML jc
            if not even_centered:
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None and jc.get(qn('w:val')) == 'center':
                        even_centered = True

            # Check small caps on text-bearing runs
            for r in para.runs:
                if r.text.strip():
                    rPr = r._element.find(qn('w:rPr'))
                    if rPr is not None:
                        sc = rPr.find(qn('w:smallCaps'))
                        if sc is not None:
                            sc_val = sc.get(qn('w:val'))
                            if sc_val is None or sc_val in ('1', 'true'):
                                even_small_caps = True
                    # Check font size (should be 9pt = sz val 18)
                    if r.font.size is not None and abs(r.font.size.pt - 9.0) < 0.5:
                        even_font_size_ok = True

        has_correct_text = 'the great adventure' in even_text.lower()
        sub_checks = [has_correct_text, even_small_caps, even_centered]
        passed = sum(sub_checks)

        if passed == 3:
            print(f"PASS: Component 3 — Even header: '{even_text}', smallCaps={even_small_caps}, centered={even_centered} (0.20 pts)")
            total_score += 0.20
        elif passed >= 1:
            partial = round(0.20 * passed / 3, 2)
            print(f"PARTIAL: Component 3 — {passed}/3 sub-checks: text='{even_text}' ({has_correct_text}), smallCaps={even_small_caps}, centered={even_centered} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Even header text='{even_text}', linked={even_linked}, smallCaps={even_small_caps}, centered={even_centered}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Odd (right) page header has STYLEREF field for chapter
    #   title in small caps, center-aligned (0.20 points)
    # In initial: "Novel Title" plain text, no field, no small caps
    # In golden: STYLEREF "Heading 1" field, small caps, centered
    # ---------------------------------------------------------------
    try:
        sec = doc.sections[0]
        odd_hdr = sec.header  # default header = odd pages when even/odd enabled

        has_field = False
        has_styleref = False
        odd_small_caps = False
        odd_centered = False

        if odd_hdr.paragraphs:
            para = odd_hdr.paragraphs[0]
            # Check for field codes in XML
            xml_str = para._element.xml
            if 'fldChar' in xml_str:
                has_field = True
            if 'STYLEREF' in xml_str:
                has_styleref = True

            # Check alignment
            pf_align = para.paragraph_format.alignment
            if pf_align is not None and pf_align == 1:  # CENTER
                odd_centered = True
            if not odd_centered:
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None and jc.get(qn('w:val')) == 'center':
                        odd_centered = True

            # Check small caps on runs
            for r in para.runs:
                rPr = r._element.find(qn('w:rPr'))
                if rPr is not None:
                    sc = rPr.find(qn('w:smallCaps'))
                    if sc is not None:
                        sc_val = sc.get(qn('w:val'))
                        if sc_val is None or sc_val in ('1', 'true'):
                            odd_small_caps = True
                            break

        # The key differentiator from initial: STYLEREF field + small caps
        # Initial has "Novel Title" as plain text, no field, no small caps
        # Centered alignment is a precondition (true in initial too), so only
        # award points if at least one task-introduced change (STYLEREF or smallCaps) is present.
        task_changes = [has_styleref, odd_small_caps]
        task_change_count = sum(task_changes)

        if has_styleref and odd_small_caps and odd_centered:
            print(f"PASS: Component 4 — Odd header: STYLEREF={has_styleref}, smallCaps={odd_small_caps}, centered={odd_centered} (0.20 pts)")
            total_score += 0.20
        elif task_change_count >= 1:
            # Only award partial for task-introduced changes, not preconditions
            partial = round(0.20 * (task_change_count + (1 if odd_centered else 0)) / 3, 2)
            print(f"PARTIAL: Component 4 — STYLEREF={has_styleref}, smallCaps={odd_small_caps}, centered={odd_centered} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Odd header: field={has_field}, STYLEREF={has_styleref}, smallCaps={odd_small_caps}, centered={odd_centered}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---------------------------------------------------------------
    # Component 5: Both headers have thin bottom border (0.10 points)
    # In initial: no bottom border on header
    # In golden: w:pBdr/w:bottom with val="single"
    # ---------------------------------------------------------------
    try:
        sec = doc.sections[0]
        borders_ok = 0

        for name, hdr in [('odd', sec.header), ('even', sec.even_page_header)]:
            if hdr.paragraphs:
                para = hdr.paragraphs[0]
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        bottom = pBdr.find(qn('w:bottom'))
                        if bottom is not None:
                            val = bottom.get(qn('w:val'))
                            if val and val != 'none':
                                borders_ok += 1
                                print(f"  {name} header: bottom border found (val={val})")

        if borders_ok == 2:
            print(f"PASS: Component 5 — Both headers have bottom border (0.10 pts)")
            total_score += 0.10
        elif borders_ok == 1:
            print(f"PARTIAL: Component 5 — Only {borders_ok}/2 headers have bottom border (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No bottom borders on headers")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ---------------------------------------------------------------
    # Component 6: Inner/outer margins set to 3.0cm / 2.0cm (0.15 points)
    # In initial: 1 inch (2.54cm) each side
    # In golden: 3.0cm inner, 2.0cm outer (with mirrored margins)
    # ---------------------------------------------------------------
    try:
        sec = doc.sections[0]
        left_cm = sec.left_margin / 914400 * 2.54
        right_cm = sec.right_margin / 914400 * 2.54

        # With mirrored margins, left_margin = inner margin, right_margin = outer margin
        inner_ok = abs(left_cm - 3.0) < 0.15
        outer_ok = abs(right_cm - 2.0) < 0.15

        if inner_ok and outer_ok:
            print(f"PASS: Component 6 — Inner={left_cm:.2f}cm, Outer={right_cm:.2f}cm (0.15 pts)")
            total_score += 0.15
        elif inner_ok or outer_ok:
            print(f"PARTIAL: Component 6 — Inner={left_cm:.2f}cm ({inner_ok}), Outer={right_cm:.2f}cm ({outer_ok}) (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 6 — Inner={left_cm:.2f}cm (expected 3.0), Outer={right_cm:.2f}cm (expected 2.0)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
