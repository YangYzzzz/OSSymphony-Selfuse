"""
Initial Setup: Employee roster document with names and departments, no tabstops
Task ID: osworld_writer_tabstop_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_tabstop_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Heading
    heading = doc.add_heading('Employee Roster', level=1)

    # 15 employees: 3-word name + 1 department word, all left-aligned, no tabstops
    employees = [
        ('John', 'Michael', 'Smith',    'Engineering'),
        ('Sarah', 'Anne', 'Brown',      'Marketing'),
        ('David', 'Lee', 'Park',        'Operations'),
        ('Emily', 'Grace', 'Johnson',   'Finance'),
        ('Robert', 'James', 'Williams', 'Engineering'),
        ('Jessica', 'Marie', 'Davis',   'Human Resources'),
        ('Michael', 'Thomas', 'Wilson', 'Marketing'),
        ('Ashley', 'Nicole', 'Moore',   'Operations'),
        ('Christopher', 'Alan', 'Taylor', 'Engineering'),
        ('Amanda', 'Rose', 'Anderson',  'Finance'),
        ('Matthew', 'Scott', 'Jackson', 'Sales'),
        ('Stephanie', 'Lynn', 'White',  'Marketing'),
        ('Daniel', 'Paul', 'Harris',    'Engineering'),
        ('Lauren', 'Elizabeth', 'Martin', 'Operations'),
        ('Kevin', 'Patrick', 'Thompson', 'Finance'),
    ]

    for first, middle, last, dept in employees:
        line = f'{first} {middle} {last} {dept}'
        para = doc.add_paragraph(line)
        # No tabstops added — all text is plain left-aligned, no tab characters
        para.paragraph_format.alignment = None  # default (left)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
