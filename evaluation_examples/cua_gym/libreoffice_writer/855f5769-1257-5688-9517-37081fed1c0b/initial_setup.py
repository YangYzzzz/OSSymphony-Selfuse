"""
Initial Setup: Banner Instructions Document + Template Image
Task ID: osworld_multi_apps_writer_to_gimp_006
Domain: multi_apps (libreoffice_writer + gimp)

Creates:
  - /home/user/Desktop/banner_instructions.docx — instructions for banner editing
  - /home/user/Desktop/banner_template.png — initial banner image to be modified
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

DESKTOP = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_to_gimp_006'
DOC_OUTPUT = f'{DESKTOP}/banner_instructions.docx'
IMG_OUTPUT = f'{DESKTOP}/banner_template.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_banner_instructions():
    """Create the banner_instructions.docx with specific editing steps."""
    doc = Document()

    # Title
    title = doc.add_heading('Banner Editing Instructions', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    doc.add_paragraph(
        'Please follow the steps below to modify the banner template image. '
        'Apply all instructions in order and save the result as banner_final.png on the Desktop.'
    )

    doc.add_paragraph('')  # blank line

    # Step 1: Resize
    doc.add_heading('Step 1: Resize the Image', level=2)
    p = doc.add_paragraph()
    p.add_run('Resize the canvas to exactly ').font.size = Pt(12)
    run_dim = p.add_run('1200 x 400 pixels')
    run_dim.bold = True
    run_dim.font.size = Pt(12)
    run_dim.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)
    p.add_run(' (width x height).').font.size = Pt(12)

    # Step 2: Background color
    doc.add_heading('Step 2: Change Background Color', level=2)
    p = doc.add_paragraph()
    p.add_run('Fill the entire background with the color ').font.size = Pt(12)
    run_color = p.add_run('#1A237E')
    run_color.bold = True
    run_color.font.size = Pt(12)
    run_color.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    p.add_run(' (dark navy blue).').font.size = Pt(12)

    # Step 3: Add text overlay
    doc.add_heading('Step 3: Add Text Overlay', level=2)
    p = doc.add_paragraph()
    p.add_run('Add the text ').font.size = Pt(12)
    run_text = p.add_run('"SUMMER SALE"')
    run_text.bold = True
    run_text.font.size = Pt(12)
    run_text.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    p.add_run(
        ' centered horizontally and vertically on the banner. '
        'Use a large, bold font (font size 72) in '
    ).font.size = Pt(12)
    run_white = p.add_run('white (#FFFFFF)')
    run_white.bold = True
    run_white.font.size = Pt(12)
    run_white.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
    p.add_run(' color.').font.size = Pt(12)

    # Step 4: Save
    doc.add_heading('Step 4: Save the Result', level=2)
    p = doc.add_paragraph()
    p.add_run('Save the modified image as ').font.size = Pt(12)
    run_save = p.add_run('banner_final.png')
    run_save.bold = True
    run_save.font.size = Pt(12)
    p.add_run(' on the Desktop. Do NOT overwrite banner_template.png.').font.size = Pt(12)

    # Footer note
    doc.add_paragraph('')
    note = doc.add_paragraph(
        'Note: The final image must be exactly 1200x400 pixels with the navy blue background '
        'and white "SUMMER SALE" text centered on the banner.'
    )
    note.runs[0].font.size = Pt(10)
    note.runs[0].italic = True

    doc.save(DOC_OUTPUT)
    print(f'Instructions document created: {DOC_OUTPUT}')


def create_banner_template():
    """Create the initial banner_template.png (800x300, light theme)."""
    width, height = 800, 300

    # Light gradient background
    img = Image.new('RGB', (width, height), color=(230, 240, 255))
    draw = ImageDraw.Draw(img)

    # Decorative background stripes
    for i in range(0, width, 40):
        draw.line([(i, 0), (i, height)], fill=(210, 225, 250), width=1)
    for j in range(0, height, 40):
        draw.line([(0, j), (width, j)], fill=(210, 225, 250), width=1)

    # Draw a decorative header bar at the top
    draw.rectangle([0, 0, width, 60], fill=(100, 149, 237))

    # Header bar text "YOUR BRAND"
    try:
        header_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 28)
        body_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 20)
        small_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 14)
    except OSError:
        header_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    # Brand name in header
    brand_text = 'YOUR BRAND'
    bbox = draw.textbbox((0, 0), brand_text, font=header_font)
    text_w = bbox[2] - bbox[0]
    draw.text(((width - text_w) // 2, 16), brand_text, fill=(255, 255, 255), font=header_font)

    # Subtitle text
    subtitle = 'Welcome to Our Store'
    bbox2 = draw.textbbox((0, 0), subtitle, font=body_font)
    sub_w = bbox2[2] - bbox2[0]
    draw.text(((width - sub_w) // 2, 100), subtitle, fill=(50, 70, 120), font=body_font)

    # Some decorative stars / sparkles
    star_positions = [(100, 160), (700, 160), (400, 200)]
    for sx, sy in star_positions:
        draw.ellipse([sx - 8, sy - 8, sx + 8, sy + 8], fill=(255, 215, 0), outline=(200, 170, 0))

    # Bottom tagline
    tagline = 'Quality Products | Best Prices | Fast Delivery'
    bbox3 = draw.textbbox((0, 0), tagline, font=small_font)
    tag_w = bbox3[2] - bbox3[0]
    draw.text(((width - tag_w) // 2, 250), tagline, fill=(80, 100, 150), font=small_font)

    # Border around image
    draw.rectangle([0, 0, width - 1, height - 1], outline=(100, 149, 237), width=3)

    img.save(IMG_OUTPUT, 'PNG')
    print(f'Banner template created: {IMG_OUTPUT}')


def main():
    os.makedirs(DESKTOP, exist_ok=True)

    create_banner_instructions()
    create_banner_template()

    # GUI-ready startup: open the instructions document in LibreOffice Writer
    # and open the template image in GIMP
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=3.0)
    launch_gui(f'gimp "{IMG_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer and GIMP with DISPLAY=:0')


main()
