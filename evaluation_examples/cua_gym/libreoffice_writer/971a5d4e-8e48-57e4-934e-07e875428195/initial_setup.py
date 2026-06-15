"""
Initial Setup: Insert a non-breaking space between Dr. and Williams in the first paragraph
Task ID: writer_txtfmt_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_txtfmt_028'
OUTPUT = f'{WORKDIR}/Desktop/medical_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # First paragraph — regular space between "Dr." and "Williams" (NOT non-breaking)
    # This is the target paragraph for the task
    para1 = doc.add_paragraph(
        "The patient was referred by Dr. Williams for a comprehensive evaluation on "
        "January 15, 2025. Initial assessment indicates moderate improvement since the "
        "previous consultation."
    )

    # Second paragraph — background context
    para2 = doc.add_paragraph(
        "The patient, a 52-year-old male, presented with chronic lower back pain "
        "originating from a workplace injury sustained in October 2023. Dr. Martinez "
        "previously noted significant inflammation in the L4-L5 region. Current "
        "medication includes ibuprofen 600 mg twice daily and physical therapy sessions "
        "three times per week."
    )

    # Third paragraph — treatment details
    para3 = doc.add_paragraph(
        "Diagnostic imaging performed on January 10, 2025 revealed mild disc "
        "degeneration but no acute herniation. The MRI results were reviewed by "
        "Dr. Nguyen, who confirmed that surgical intervention is not required at "
        "this stage. Conservative management will continue for an additional six weeks."
    )

    # Fourth paragraph — follow-up instructions
    para4 = doc.add_paragraph(
        "Follow-up appointment is scheduled for February 28, 2025. The patient is "
        "advised to maintain a daily exercise regimen as outlined in the physiotherapy "
        "plan. Any significant worsening of symptoms should be reported immediately to "
        "the primary care team."
    )

    # Fifth paragraph — signature block
    para5 = doc.add_paragraph(
        "This report was prepared by the Department of Physical Medicine and "
        "Rehabilitation, City General Hospital. For inquiries, please contact the "
        "Medical Records Office at extension 4217."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
