"""
Reward Script: Insert cross-reference to bookmark 'table_results' displaying 'Table 3'
Task ID: writer_bs_019
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): REF field referencing 'table_results' bookmark exists
  Component 2 (0.3): Field display text contains 'Table 3'
  Component 3 (0.3): Cross-reference is located in paragraph containing 'as shown in'
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_019'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify cross-reference insertion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: REF field referencing 'table_results' bookmark exists (0.4 points)
    # A cross-reference to a bookmark is implemented as a REF field with instrText containing the bookmark name.
    try:
        instr_texts = doc.element.findall('.//w:instrText', ns)
        ref_field_found = False
        for it in instr_texts:
            if it.text and 'REF' in it.text and 'table_results' in it.text:
                ref_field_found = True
                break

        if ref_field_found:
            print(f"PASS: Component 1 -- REF field referencing 'table_results' found (0.4 pts)")
            total_score += 0.4
        else:
            instr_values = [it.text for it in instr_texts if it.text]
            print(f"FAIL: Component 1 -- No REF field referencing 'table_results'. Found instrTexts: {instr_values}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Field display text contains 'Table 3' (0.3 points)
    # Between fldChar separate and end, there should be a run with text 'Table 3'
    try:
        field_display_text = ""
        # Walk through all paragraphs and find the field structure
        for para in doc.paragraphs:
            runs = para._element.findall('.//w:r', ns)
            in_field = False
            past_separate = False
            current_field_is_ref_table = False

            for run in runs:
                # Check for instrText with REF table_results
                instr = run.find('w:instrText', ns)
                if instr is not None and instr.text and 'REF' in instr.text and 'table_results' in instr.text:
                    current_field_is_ref_table = True

                # Check for fldChar
                fld_char = run.find('w:fldChar', ns)
                if fld_char is not None:
                    fld_type = fld_char.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', '')
                    if fld_type == 'begin':
                        in_field = True
                        past_separate = False
                        current_field_is_ref_table = False
                    elif fld_type == 'separate':
                        past_separate = True
                    elif fld_type == 'end':
                        in_field = False
                        past_separate = False
                        current_field_is_ref_table = False

                # Collect display text after 'separate' in the REF table_results field
                if in_field and past_separate and current_field_is_ref_table:
                    t_elem = run.find('w:t', ns)
                    if t_elem is not None and t_elem.text:
                        field_display_text += t_elem.text

        if 'Table 3' in field_display_text:
            print(f"PASS: Component 2 -- Field displays 'Table 3' (display text: '{field_display_text.strip()}') (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 -- Expected 'Table 3' in field display, found: '{field_display_text.strip()}'")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Cross-reference is in paragraph containing 'as shown in' (0.3 points)
    # The task says the cursor is after 'as shown in ' on page 3, so the REF field should be in that paragraph.
    try:
        ref_in_correct_para = False
        for para in doc.paragraphs:
            para_text = para.text.lower()
            # Check if this paragraph has both 'as shown in' text AND a REF field
            if 'as shown in' in para_text:
                # Check if this paragraph contains a REF table_results field
                instr_elems = para._element.findall('.//w:instrText', ns)
                for ie in instr_elems:
                    if ie.text and 'REF' in ie.text and 'table_results' in ie.text:
                        ref_in_correct_para = True
                        break
            if ref_in_correct_para:
                break

        if ref_in_correct_para:
            print(f"PASS: Component 3 -- Cross-reference located in paragraph with 'as shown in' (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 -- Cross-reference not found in paragraph containing 'as shown in'")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
