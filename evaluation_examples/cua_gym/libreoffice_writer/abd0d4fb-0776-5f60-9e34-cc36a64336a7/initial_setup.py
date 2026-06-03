"""
Initial Setup: Insert cross-reference to bookmark 'table_results'
Task ID: writer_bs_019
Domain: libreoffice_writer

Creates a multi-page research document with:
- Several pages of academic content
- Multiple tables throughout (Table 1, Table 2, Table 3)
- A bookmark 'table_results' at the "Table 3: Experimental Results" caption (page ~8)
- Text on page ~3 ending with "as shown in " where the cross-reference should go
- NO cross-reference field (that's the task for the agent)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark spanning the entire paragraph text."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._element

    # Create bookmark start
    bm_start = tag.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): '1',
        qn('w:name'): bookmark_name,
    })
    tag.addprevious(bm_start)

    # Create bookmark end (after last run)
    last_run = paragraph.runs[-1]._element
    bm_end = last_run.makeelement(qn('w:bookmarkEnd'), {
        qn('w:id'): '1',
    })
    last_run.addnext(bm_end)


def add_filler_paragraphs(doc, count, topic_text):
    """Add realistic filler paragraphs to build page length."""
    for text in topic_text[:count]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # =================================================================
    # PAGE 1: Title & Abstract
    # =================================================================
    title = doc.add_heading('Comparative Analysis of Machine Learning Approaches for Climate Prediction Models', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Dr. Elena Vasquez, Prof. Rajesh Patel, Dr. Mei-Lin Chang')
    run.font.size = Pt(11)
    run.font.italic = True

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil.add_run('Department of Atmospheric Sciences, Pacific Research Institute')
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    abs_heading = doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This study presents a comprehensive evaluation of three machine learning frameworks '
        'applied to regional climate prediction across the Pacific Northwest corridor. We compare '
        'gradient boosting methods, recurrent neural architectures, and transformer-based models '
        'using a dataset spanning 42 years of meteorological observations from 156 monitoring '
        'stations. Our findings indicate that hybrid transformer-LSTM models achieve superior '
        'accuracy for precipitation forecasting while gradient boosting methods remain competitive '
        'for temperature anomaly detection. The implications for operational weather forecasting '
        'systems are discussed, along with recommendations for model deployment at scale.'
    )

    doc.add_paragraph(
        'Keywords: climate modeling, machine learning, precipitation forecasting, transformer '
        'networks, gradient boosting, LSTM, regional weather prediction, Pacific Northwest'
    )

    # =================================================================
    # PAGE 2: Introduction
    # =================================================================
    doc.add_heading('1. Introduction', level=1)
    intro_texts = [
        'Climate prediction remains one of the most challenging problems in atmospheric science, '
        'requiring the integration of complex physical processes with statistical modeling frameworks. '
        'Traditional numerical weather prediction (NWP) models, while physically grounded, often '
        'struggle with capturing mesoscale phenomena and local topographic effects that significantly '
        'influence regional weather patterns.',

        'The emergence of machine learning (ML) techniques has opened new avenues for improving '
        'prediction accuracy, particularly for variables such as precipitation amount, extreme weather '
        'events, and temperature anomalies at sub-regional scales. Several recent studies have '
        'demonstrated that ML-based post-processing of NWP output can reduce forecast errors by '
        '15-30% compared to raw model output (Harrison et al., 2024; Nakamura & Singh, 2023).',

        'However, the relative performance of different ML architectures for specific climate '
        'variables remains poorly understood. Most comparative studies focus on single prediction '
        'targets or limited geographic regions, making it difficult to draw generalizable conclusions '
        'about which methods are best suited for operational deployment.',

        'In this paper, we address these gaps by conducting a systematic comparison of three '
        'distinct ML paradigms: (1) gradient boosting machines (XGBoost, LightGBM), (2) recurrent '
        'neural networks (LSTM, GRU), and (3) transformer-based architectures (temporal fusion '
        'transformers). Our evaluation spans multiple prediction horizons (24h, 72h, 168h) across '
        'the Pacific Northwest region, which presents diverse topographic and climatic challenges '
        'ranging from coastal marine influences to inland continental conditions.',
    ]
    add_filler_paragraphs(doc, len(intro_texts), intro_texts)

    # =================================================================
    # PAGE 3: Data & Methods (with "as shown in " text)
    # =================================================================
    doc.add_heading('2. Data and Methods', level=1)

    doc.add_heading('2.1 Study Area and Data Sources', level=2)
    data_texts = [
        'Our study area encompasses the Pacific Northwest corridor, defined as the region between '
        '42°N to 49°N latitude and 116°W to 125°W longitude. This area includes portions of '
        'Washington, Oregon, and northern California, featuring significant topographic variation '
        'from sea level to elevations exceeding 4,300 meters at Mount Rainier.',

        'Observational data were obtained from 156 automated weather stations operated by the '
        'National Weather Service (NWS), the Remote Automatic Weather Stations (RAWS) network, '
        'and the MesoWest cooperative network. The dataset covers the period from January 1982 '
        'through December 2023, providing 42 years of continuous hourly measurements.',

        'Variables collected include surface temperature (2m), precipitation accumulation, relative '
        'humidity, wind speed and direction, barometric pressure, and solar radiation. Quality '
        'control procedures followed the methods described by Durre et al. (2010), with additional '
        'spatial consistency checks implemented using neighboring station cross-validation.',
    ]
    add_filler_paragraphs(doc, len(data_texts), data_texts)

    # Table 1: Station Summary
    doc.add_paragraph()
    t1_caption = doc.add_paragraph()
    run = t1_caption.add_run('Table 1: Summary of Weather Station Networks')
    run.bold = True
    t1_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Table Grid'
    headers = ['Network', 'Stations', 'Elevation Range (m)', 'Record Period']
    for i, h in enumerate(headers):
        table1.cell(0, i).text = h
        for run in table1.cell(0, i).paragraphs[0].runs:
            run.bold = True
    data1 = [
        ['NWS ASOS', '47', '0 – 1,245', '1982 – 2023'],
        ['RAWS', '62', '120 – 3,890', '1988 – 2023'],
        ['MesoWest', '38', '5 – 2,100', '1995 – 2023'],
        ['University Network', '9', '45 – 1,560', '2001 – 2023'],
    ]
    for r, row_data in enumerate(data1, 1):
        for c, val in enumerate(row_data):
            table1.cell(r, c).text = val

    doc.add_paragraph()

    doc.add_heading('2.2 Feature Engineering', level=2)
    feature_texts = [
        'Feature extraction was performed at multiple temporal scales to capture both synoptic-scale '
        'forcing and local diurnal cycles. For each station, we computed rolling statistics (mean, '
        'standard deviation, min, max) over windows of 6, 12, 24, and 48 hours for all primary '
        'variables.',

        'Derived features included temperature lapse rate estimates (computed from elevation '
        'differences between neighboring stations), moisture flux convergence, and the Heidke '
        'Skill Score for categorical precipitation thresholds. Temporal encoding features captured '
        'day-of-year, hour-of-day, and lunar phase using sinusoidal transformations.',
    ]
    add_filler_paragraphs(doc, len(feature_texts), feature_texts)

    # KEY PARAGRAPH: "as shown in " — this is where the agent should insert the cross-ref
    ref_para = doc.add_paragraph(
        'The performance metrics for all three model families across each prediction horizon are '
        'summarized in the experimental results section, as shown in '
    )
    ref_para.paragraph_format.space_after = Pt(6)

    # Continue after the reference placeholder
    doc.add_paragraph(
        'which provides a detailed breakdown by geographic sub-region and weather variable. '
        'These results highlight the strengths and weaknesses of each approach under varying '
        'atmospheric conditions.'
    )

    # =================================================================
    # PAGE 4-5: Model Architectures
    # =================================================================
    doc.add_heading('2.3 Model Architectures', level=2)

    doc.add_heading('2.3.1 Gradient Boosting Methods', level=3)
    gb_texts = [
        'We implemented two gradient boosting frameworks: XGBoost (Chen & Guestrin, 2016) and '
        'LightGBM (Ke et al., 2017). Both models were trained with early stopping based on '
        'validation RMSE, using a 70/15/15 train/validation/test split stratified by year to '
        'prevent temporal leakage.',

        'Hyperparameter optimization was conducted using Bayesian optimization with Tree-structured '
        'Parzen Estimators (TPE) over 200 trials for each configuration. Key parameters optimized '
        'included learning rate (0.01–0.3), maximum depth (3–12), minimum child weight (1–10), '
        'subsample ratio (0.5–1.0), and column sample ratio (0.3–1.0). Feature importance was '
        'assessed using SHAP values to ensure physical interpretability of the model predictions.',

        'For precipitation forecasting specifically, we employed a two-stage approach: first '
        'classifying precipitation occurrence (binary), then predicting amount conditional on '
        'occurrence. This decomposition has been shown to improve skill for the inherently '
        'discontinuous precipitation variable (Scheuerer & Hamill, 2015).',
    ]
    add_filler_paragraphs(doc, len(gb_texts), gb_texts)

    doc.add_heading('2.3.2 Recurrent Neural Networks', level=3)
    rnn_texts = [
        'Long Short-Term Memory (LSTM) networks were configured with 2 hidden layers of 256 units '
        'each, followed by a fully connected output layer. Input sequences of 168 time steps '
        '(7 days of hourly data) were used to predict the target variable at lead times of 24, '
        '72, and 168 hours.',

        'Training employed Adam optimizer with cosine annealing learning rate schedule, starting '
        'at 1e-3 and decaying to 1e-6 over 100 epochs. Dropout of 0.2 was applied between LSTM '
        'layers, and gradient clipping was set at 1.0 to prevent exploding gradients. Batch size '
        'was 256 samples, with sequences shuffled at the epoch level while maintaining temporal '
        'ordering within each sequence.',

        'We also evaluated Gated Recurrent Unit (GRU) variants with identical architecture to '
        'assess whether the simpler gating mechanism provides comparable performance with reduced '
        'computational cost. All neural network models were implemented in PyTorch 2.1 and trained '
        'on NVIDIA A100 GPUs with mixed-precision (FP16) acceleration.',
    ]
    add_filler_paragraphs(doc, len(rnn_texts), rnn_texts)

    doc.add_heading('2.3.3 Transformer-Based Models', level=3)
    tf_texts = [
        'The Temporal Fusion Transformer (TFT) architecture (Lim et al., 2021) was adapted for '
        'our multi-horizon prediction task. The model employs variable selection networks to '
        'automatically identify relevant input features, gated residual connections for efficient '
        'information flow, and multi-head attention with 4 heads operating over the temporal '
        'dimension.',

        'A key modification to the standard TFT was the inclusion of spatial attention across '
        'neighboring stations, allowing the model to learn geographic dependencies without '
        'explicit spatial feature engineering. Station embeddings of dimension 32 were learned '
        'jointly with the temporal model parameters.',

        'The hybrid Transformer-LSTM variant replaced the TFT encoder with an LSTM-based sequence '
        'processor while retaining the multi-head temporal attention decoder. This design aims to '
        'combine the sequential modeling strengths of LSTMs with the long-range dependency capture '
        'of attention mechanisms.',
    ]
    add_filler_paragraphs(doc, len(tf_texts), tf_texts)

    # =================================================================
    # PAGE 5-6: Table 2 and more methods
    # =================================================================
    doc.add_heading('2.4 Evaluation Metrics', level=2)
    eval_texts = [
        'Model performance was assessed using a comprehensive suite of metrics appropriate for '
        'each prediction target. For continuous variables (temperature, pressure), we report Root '
        'Mean Square Error (RMSE), Mean Absolute Error (MAE), and the coefficient of determination '
        '(R²). For precipitation amount, we additionally report the Continuous Ranked Probability '
        'Score (CRPS) to assess probabilistic calibration.',

        'Categorical metrics for precipitation occurrence include the Brier Skill Score (BSS), '
        'area under the ROC curve (AUC-ROC), and the reliability diagram analysis. Extreme event '
        'detection is evaluated using the Extreme Dependency Score (EDS) for events above the '
        '95th percentile threshold.',
    ]
    add_filler_paragraphs(doc, len(eval_texts), eval_texts)

    # Table 2: Hyperparameters
    doc.add_paragraph()
    t2_caption = doc.add_paragraph()
    run = t2_caption.add_run('Table 2: Optimal Hyperparameters by Model Family')
    run.bold = True
    t2_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Table Grid'
    t2_headers = ['Parameter', 'Gradient Boosting', 'LSTM/GRU']
    for i, h in enumerate(t2_headers):
        table2.cell(0, i).text = h
        for run in table2.cell(0, i).paragraphs[0].runs:
            run.bold = True
    t2_data = [
        ['Learning Rate', '0.045', '0.001'],
        ['Hidden Layers/Depth', '8', '2 x 256'],
        ['Regularization', 'L2 = 0.1', 'Dropout 0.2'],
        ['Batch Size', '1024', '256'],
        ['Training Epochs', '450 (early stop)', '100'],
        ['Feature Count', '87', '87'],
    ]
    for r, row_data in enumerate(t2_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_paragraph()

    # =================================================================
    # PAGE 6-7: Results discussion
    # =================================================================
    doc.add_heading('3. Results', level=1)

    doc.add_heading('3.1 Temperature Prediction', level=2)
    temp_texts = [
        'Temperature predictions showed the smallest inter-model differences among all variables '
        'tested. At the 24-hour lead time, all model families achieved RMSE values below 1.8°C, '
        'with the Transformer-LSTM hybrid producing the lowest error (1.42°C). Gradient boosting '
        'methods performed comparably at 1.51°C RMSE, while standalone LSTM models showed slightly '
        'higher errors at 1.67°C.',

        'Performance degradation with increasing lead time was most pronounced for the LSTM '
        'models, which showed a 68% increase in RMSE from 24h to 168h forecasts. In contrast, '
        'gradient boosting methods degraded by only 41%, suggesting better utilization of '
        'climatological features that remain relevant at longer horizons.',

        'Spatial analysis revealed that coastal stations showed consistently lower prediction '
        'errors across all models, attributable to the moderating influence of ocean temperatures. '
        'Inland stations at higher elevations exhibited the largest inter-model performance gaps, '
        'with the Transformer-LSTM showing particular advantage in capturing orographic effects.',
    ]
    add_filler_paragraphs(doc, len(temp_texts), temp_texts)

    doc.add_heading('3.2 Precipitation Forecasting', level=2)
    precip_texts = [
        'Precipitation prediction presented the greatest challenge and the largest performance '
        'differences between model families. The two-stage gradient boosting approach achieved '
        'the highest Brier Skill Score (0.42) for precipitation occurrence at 24h lead time, '
        'while the Transformer-LSTM produced the best continuous predictions (CRPS = 0.87 mm) '
        'for precipitation amount.',

        'For extreme precipitation events (above 95th percentile), the transformer-based models '
        'showed clear superiority with an EDS of 0.73 compared to 0.58 for gradient boosting '
        'and 0.51 for LSTM. This advantage is attributed to the attention mechanism capacity to '
        'identify rare but influential atmospheric patterns in the training data.',

        'Seasonal analysis showed that winter precipitation (dominated by synoptic-scale frontal '
        'systems) was better predicted by all models compared to summer convective precipitation. '
        'The performance gap was smallest for the Transformer-LSTM, suggesting better adaptability '
        'to varying precipitation mechanisms.',
    ]
    add_filler_paragraphs(doc, len(precip_texts), precip_texts)

    # =================================================================
    # PAGE 8: Table 3 with bookmark
    # =================================================================
    doc.add_heading('3.3 Comprehensive Model Comparison', level=2)

    doc.add_paragraph(
        'The following table presents the comprehensive results across all model architectures, '
        'prediction variables, and forecast horizons. These metrics represent averages over the '
        'held-out test period (2020-2023), with confidence intervals computed via block bootstrap '
        'resampling with a block length of 30 days.'
    )

    doc.add_paragraph()

    # Table 3 caption with bookmark
    t3_caption = doc.add_paragraph()
    run = t3_caption.add_run('Table 3: Experimental Results')
    run.bold = True
    t3_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Add bookmark 'table_results' to this caption paragraph
    add_bookmark(t3_caption, 'table_results')

    table3 = doc.add_table(rows=8, cols=5)
    table3.style = 'Table Grid'
    t3_headers = ['Model', 'Temp RMSE (°C)', 'Precip CRPS (mm)', 'Wind MAE (m/s)', 'Overall R²']
    for i, h in enumerate(t3_headers):
        table3.cell(0, i).text = h
        for run in table3.cell(0, i).paragraphs[0].runs:
            run.bold = True
    t3_data = [
        ['XGBoost', '1.51', '1.23', '1.89', '0.847'],
        ['LightGBM', '1.48', '1.19', '1.85', '0.853'],
        ['LSTM', '1.67', '1.41', '2.04', '0.821'],
        ['GRU', '1.63', '1.38', '1.98', '0.829'],
        ['TFT', '1.45', '1.08', '1.76', '0.868'],
        ['Transformer-LSTM', '1.42', '0.97', '1.71', '0.882'],
        ['NWP Baseline', '2.34', '2.15', '2.67', '0.734'],
    ]
    for r, row_data in enumerate(t3_data, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    doc.add_paragraph()

    # =================================================================
    # PAGE 8-9: Discussion
    # =================================================================
    doc.add_heading('4. Discussion', level=1)
    disc_texts = [
        'Our results demonstrate that no single model family dominates across all prediction '
        'targets and horizons. The Transformer-LSTM hybrid shows the most consistent performance, '
        'achieving top or near-top scores for temperature, precipitation, and wind prediction. '
        'However, gradient boosting methods remain highly competitive, particularly for shorter '
        'lead times and for precipitation occurrence classification.',

        'The computational cost analysis reveals important practical considerations. Gradient '
        'boosting models require approximately 2 hours of training on a single CPU core, while '
        'the Transformer-LSTM requires 18 hours on an A100 GPU. For operational deployment where '
        'models must be retrained frequently, this 9x cost difference may favor the simpler '
        'methods despite their slightly lower accuracy.',

        'An important finding is the geographic dependency of model performance. The attention '
        'mechanism in transformer models appears particularly beneficial in areas with complex '
        'topography, where long-range spatial dependencies carry significant predictive information. '
        'This suggests that operational systems might benefit from a geographically-weighted ensemble '
        'approach that leverages different model strengths in different terrain types.',
    ]
    add_filler_paragraphs(doc, len(disc_texts), disc_texts)

    doc.add_heading('5. Conclusions', level=1)
    doc.add_paragraph(
        'This study provides a systematic comparison of machine learning approaches for regional '
        'climate prediction in the Pacific Northwest. Our key findings are: (1) Transformer-LSTM '
        'hybrids achieve the best overall accuracy, particularly for extreme events; (2) gradient '
        'boosting methods offer the best accuracy-to-cost ratio for operational applications; '
        '(3) model performance is strongly geography-dependent, suggesting the need for adaptive '
        'ensemble strategies. Future work will extend this analysis to additional climate variables '
        'and explore transfer learning approaches for data-sparse regions.'
    )

    # =================================================================
    # Save
    # =================================================================
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
