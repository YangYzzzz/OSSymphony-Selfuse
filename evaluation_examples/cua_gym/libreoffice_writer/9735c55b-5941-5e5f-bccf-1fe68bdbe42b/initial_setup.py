"""
Initial Setup: Create a document with a legal disclaimer paragraph for AutoText creation
Task ID: writer_frd_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_051'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

DISCLAIMER_TEXT = (
    'DISCLAIMER: This document is provided for informational purposes only '
    'and does not constitute legal advice. The information contained herein '
    'is subject to change without notice. No liability is assumed for any '
    'errors or omissions.'
)


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Contract Services Agreement', level=1)

    # Introductory paragraph
    p1 = doc.add_paragraph()
    run1 = p1.add_run(
        'This Agreement is entered into as of March 15, 2025, by and between '
        'Meridian Technology Solutions, Inc. ("Provider") and Apex Financial '
        'Group, LLC ("Client"). The Provider agrees to deliver consulting '
        'services as outlined in Exhibit A attached hereto.'
    )
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)

    # Scope of services paragraph
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Scope of Services')
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        'The Provider shall furnish technical advisory services including but '
        'not limited to system architecture review, security audit consultation, '
        'and infrastructure optimization recommendations. All deliverables shall '
        'be submitted within 30 business days of the engagement start date.'
    )
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    # Payment terms
    p4 = doc.add_paragraph()
    run4 = p4.add_run('Payment Terms')
    run4.bold = True
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(12)

    p5 = doc.add_paragraph()
    run5 = p5.add_run(
        'Client agrees to compensate Provider at a rate of $275.00 per hour for '
        'all services rendered. Invoices shall be submitted on a bi-weekly basis '
        'and are payable within 30 days of receipt. Late payments shall incur '
        'interest at a rate of 1.5% per month.'
    )
    run5.font.name = 'Times New Roman'
    run5.font.size = Pt(12)

    # The disclaimer paragraph that the user wants to save as AutoText
    p6 = doc.add_paragraph()
    run6 = p6.add_run(DISCLAIMER_TEXT)
    run6.font.name = 'Times New Roman'
    run6.font.size = Pt(11)
    run6.italic = True
    p6.paragraph_format.space_before = Pt(12)

    # Signature block
    p7 = doc.add_paragraph()
    p7.paragraph_format.space_before = Pt(24)
    run7 = p7.add_run('Authorized Signature: ________________________')
    run7.font.name = 'Times New Roman'
    run7.font.size = Pt(12)

    p8 = doc.add_paragraph()
    run8 = p8.add_run('Date: ________________________')
    run8.font.name = 'Times New Roman'
    run8.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
