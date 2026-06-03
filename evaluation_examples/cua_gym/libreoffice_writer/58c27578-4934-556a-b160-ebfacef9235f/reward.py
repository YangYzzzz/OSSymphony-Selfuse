"""
Reward Script: Build a complete project proposal document
Task ID: writer_pd_045
Domain: libreoffice_writer
Scoring:
  C1: Cover page (centered title, client, date, logo placeholder) — 0.15
  C2: Table of Contents section with listings — 0.10
  C3: Executive Summary section — 0.10
  C4: Scope of Work with 3-level numbered outline — 0.20
  C5: Gantt chart placeholder table (phases + timeline columns) — 0.15
  C6: Pricing table with line items and totals — 0.10
  C7: Terms section with 5 numbered clauses — 0.10
  C8: Heading color #1B4F72 — 0.05
  C9: Body text color #2C3E50 — 0.05
"""

import os
import re
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_045'


def color_distance(c1, c2):
    """Euclidean RGB distance between two RGBColor objects or tuples."""
    r1, g1, b1 = (c1[0], c1[1], c1[2]) if not isinstance(c1, RGBColor) else (c1[0], c1[1], c1[2])
    r2, g2, b2 = (c2[0], c2[1], c2[2]) if not isinstance(c2, RGBColor) else (c2[0], c2[1], c2[2])
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gate: document must have meaningful content (not blank)
    if len(doc.paragraphs) < 10:
        print(f"FAIL: Document has only {len(doc.paragraphs)} paragraphs — too few for a proposal")
        print("REWARD: 0.0")
        return 0.0

    # =========================================================================
    # Component 1: Cover page — centered title, client name, date, logo (0.15)
    # =========================================================================
    try:
        # Look in first ~16 paragraphs for cover page elements
        cover_paras = doc.paragraphs[:17]
        cover_texts = [p.text.strip().lower() for p in cover_paras]
        cover_full = " ".join(cover_texts)

        has_title = False
        has_client = False
        has_date = False
        has_logo = False
        centered_count = 0

        for p in cover_paras:
            t = p.text.strip().lower()
            if not t:
                continue

            # Check centering
            if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                centered_count += 1

            # Title: a prominent proposal-related heading
            if any(kw in t for kw in ['proposal', 'initiative', 'transformation', 'project']):
                if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    has_title = True

            # Client name
            if any(kw in t for kw in ['prepared for', 'client', 'meridian', 'technologies']):
                has_client = True

            # Date
            if any(kw in t for kw in ['date', '2026', '2025', 'march', 'april']):
                has_date = True

            # Logo placeholder
            if any(kw in t for kw in ['logo', 'company logo', '[logo']):
                has_logo = True

        cover_sub = 0.0
        if has_title:
            cover_sub += 0.04
        if has_client:
            cover_sub += 0.04
        if has_date:
            cover_sub += 0.03
        if has_logo:
            cover_sub += 0.02
        if centered_count >= 3:
            cover_sub += 0.02

        if cover_sub > 0:
            print(f"PASS: Component 1 — Cover page (title={has_title}, client={has_client}, "
                  f"date={has_date}, logo={has_logo}, centered={centered_count}) ({cover_sub:.2f} pts)")
            total_score += cover_sub
        else:
            print("FAIL: Component 1 — No cover page elements found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Table of Contents section (0.10)
    # =========================================================================
    try:
        toc_found = False
        toc_entries = 0

        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip().lower()
            if 'table of contents' in t:
                toc_found = True
                # Count TOC entries (lines with dots or page numbers following)
                for j in range(i + 1, min(i + 15, len(doc.paragraphs))):
                    entry = doc.paragraphs[j].text.strip()
                    if entry and (re.search(r'\.{3,}', entry) or re.search(r'\d+$', entry)):
                        toc_entries += 1
                    # Stop at next heading
                    if doc.paragraphs[j].style.name.startswith('Heading') and j > i + 1:
                        break
                break

        if toc_found and toc_entries >= 3:
            print(f"PASS: Component 2 — TOC found with {toc_entries} entries (0.10 pts)")
            total_score += 0.10
        elif toc_found:
            print(f"PARTIAL: Component 2 — TOC heading found but only {toc_entries} entries (0.05 pts)")
            total_score += 0.05
        else:
            print("FAIL: Component 2 — No Table of Contents found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: Executive Summary section (0.10)
    # =========================================================================
    try:
        exec_found = False
        exec_body_count = 0

        for i, p in enumerate(doc.paragraphs):
            if 'executive summary' in p.text.strip().lower():
                exec_found = True
                # Count body paragraphs after the heading
                for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                    pj = doc.paragraphs[j]
                    if pj.style.name.startswith('Heading'):
                        break
                    if pj.text.strip() and len(pj.text.strip()) > 30:
                        exec_body_count += 1
                break

        if exec_found and exec_body_count >= 2:
            print(f"PASS: Component 3 — Executive Summary with {exec_body_count} body paragraphs (0.10 pts)")
            total_score += 0.10
        elif exec_found:
            print(f"PARTIAL: Component 3 — Executive Summary heading found, {exec_body_count} body paras (0.05 pts)")
            total_score += 0.05
        else:
            print("FAIL: Component 3 — No Executive Summary section found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================================
    # Component 4: Scope of Work with 3-level numbered outline (0.20)
    # =========================================================================
    try:
        scope_found = False
        level1 = 0
        level2 = 0
        level3 = 0

        in_scope = False
        for p in doc.paragraphs:
            t = p.text.strip().lower()
            # Only trigger on the Heading paragraph, not on TOC entries
            if 'scope of work' in t and p.style.name.startswith('Heading'):
                scope_found = True
                in_scope = True
                continue
            # Stop at a non-scope major heading (Heading 1 that isn't a sub-scope heading)
            if in_scope and p.style.name == 'Heading 1' and 'scope' not in t.lower():
                break
            if in_scope:
                raw = p.text.strip()
                # Level 3: indented X.X.X. pattern (may have leading spaces)
                stripped = raw.lstrip()
                if re.match(r'^\d+\.\d+\.\d+\.?\s', stripped):
                    level3 += 1
                # Level 2: indented X.X. pattern (may have leading spaces)
                elif re.match(r'^\d+\.\d+\.?\s', stripped):
                    level2 += 1
                # Level 1: top-level N. pattern (not deeply indented)
                elif re.match(r'^\d+\.\s', stripped):
                    level1 += 1

        scope_sub = 0.0
        if scope_found:
            scope_sub += 0.05
        if level1 >= 2:
            scope_sub += 0.05
        if level2 >= 3:
            scope_sub += 0.05
        if level3 >= 3:
            scope_sub += 0.05

        if scope_sub > 0:
            print(f"PASS: Component 4 — Scope of Work (found={scope_found}, "
                  f"L1={level1}, L2={level2}, L3={level3}) ({scope_sub:.2f} pts)")
            total_score += scope_sub
        else:
            print("FAIL: Component 4 — No Scope of Work with numbered outline found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =========================================================================
    # Component 5: Gantt chart placeholder table (0.15)
    # =========================================================================
    try:
        gantt_table = None
        # Look for a table that resembles a Gantt chart (phases in rows, time in columns)
        for table in doc.tables:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            # Check if headers contain time-related columns (Q1, Month, etc.)
            has_time_cols = any(
                any(kw in cell for kw in ['q1', 'q2', 'q3', 'q4', 'month', 'jan', 'feb', 'mar',
                                          '2026', '2027', 'week'])
                for cell in header_cells
            )
            has_phase_col = any(
                any(kw in cell for kw in ['phase', 'project', 'task', 'activity', 'deliverable',
                                          'milestone'])
                for cell in header_cells
            )
            if has_time_cols and (has_phase_col or len(table.columns) >= 5):
                gantt_table = table
                break

        if gantt_table is not None:
            num_rows = len(gantt_table.rows)
            num_cols = len(gantt_table.columns)
            # Check for phase names in first column (rows after header)
            phase_names = [gantt_table.cell(r, 0).text.strip()
                           for r in range(1, num_rows) if gantt_table.cell(r, 0).text.strip()]

            gantt_sub = 0.0
            if num_cols >= 4:
                gantt_sub += 0.05
            if num_rows >= 4:
                gantt_sub += 0.05
            if len(phase_names) >= 3:
                gantt_sub += 0.05

            print(f"PASS: Component 5 — Gantt table ({num_rows}x{num_cols}), "
                  f"phases={phase_names[:4]} ({gantt_sub:.2f} pts)")
            total_score += gantt_sub
        else:
            print("FAIL: Component 5 — No Gantt chart table found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # =========================================================================
    # Component 6: Pricing table with line items and totals (0.10)
    # =========================================================================
    try:
        pricing_table = None
        for table in doc.tables:
            all_text = " ".join(c.text.strip().lower() for row in table.rows for c in row.cells)
            if any(kw in all_text for kw in ['total', 'price', 'rate', 'cost', 'budget', 'investment']):
                header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
                if any(kw in " ".join(header_cells) for kw in ['total', 'rate', 'price', 'cost',
                                                                'amount', 'deliverable', 'item']):
                    pricing_table = table
                    break

        if pricing_table is not None:
            num_rows = len(pricing_table.rows)
            # Check for total row
            last_rows_text = " ".join(
                c.text.strip().lower()
                for row in pricing_table.rows[-3:]
                for c in row.cells
            )
            has_total = 'total' in last_rows_text

            # Check for dollar amounts
            dollar_count = 0
            for row in pricing_table.rows[1:]:
                for cell in row.cells:
                    if '$' in cell.text or re.search(r'\d{1,3}(,\d{3})+', cell.text):
                        dollar_count += 1
                        break

            pricing_sub = 0.0
            if num_rows >= 5 and dollar_count >= 3:
                pricing_sub += 0.05
            if has_total:
                pricing_sub += 0.05

            print(f"PASS: Component 6 — Pricing table ({num_rows} rows, "
                  f"{dollar_count} rows with amounts, total={has_total}) ({pricing_sub:.2f} pts)")
            total_score += pricing_sub
        else:
            print("FAIL: Component 6 — No pricing table found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # =========================================================================
    # Component 7: Terms section with 5 numbered clauses (0.10)
    # =========================================================================
    try:
        terms_found = False
        clause_count = 0

        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip().lower()
            if 'terms' in t and (p.style.name.startswith('Heading') or 'condition' in t):
                terms_found = True
                # Count numbered clauses after the heading
                for j in range(i + 1, len(doc.paragraphs)):
                    pj = doc.paragraphs[j]
                    raw = pj.text.strip()
                    # Numbered clause title: "N. SomeTitle"
                    if re.match(r'^\d+\.\s+[A-Z]', raw):
                        clause_count += 1
                break

        if terms_found and clause_count >= 5:
            print(f"PASS: Component 7 — Terms section with {clause_count} clauses (0.10 pts)")
            total_score += 0.10
        elif terms_found and clause_count >= 3:
            print(f"PARTIAL: Component 7 — Terms found with {clause_count} clauses (0.05 pts)")
            total_score += 0.05
        elif terms_found:
            print(f"PARTIAL: Component 7 — Terms heading found but only {clause_count} clauses (0.03 pts)")
            total_score += 0.03
        else:
            print("FAIL: Component 7 — No Terms section found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # =========================================================================
    # Component 8: Heading color #1B4F72 (0.05)
    # =========================================================================
    try:
        target_heading_color = RGBColor(0x1B, 0x4F, 0x72)
        heading_paras = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        matching_headings = 0
        total_headings = 0

        for p in heading_paras:
            for run in p.runs:
                if run.text.strip():
                    total_headings += 1
                    if run.font.color and run.font.color.rgb:
                        dist = color_distance(run.font.color.rgb, target_heading_color)
                        if dist < 30:
                            matching_headings += 1
                    break  # check first run per heading

        if total_headings > 0 and matching_headings / total_headings >= 0.7:
            print(f"PASS: Component 8 — Heading color #1B4F72 "
                  f"({matching_headings}/{total_headings} match) (0.05 pts)")
            total_score += 0.05
        elif matching_headings > 0:
            print(f"PARTIAL: Component 8 — Some headings match "
                  f"({matching_headings}/{total_headings}) (0.02 pts)")
            total_score += 0.02
        else:
            print(f"FAIL: Component 8 — No headings with color #1B4F72 "
                  f"(checked {total_headings} headings)")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    # =========================================================================
    # Component 9: Body text color #2C3E50 (0.05)
    # =========================================================================
    try:
        target_body_color = RGBColor(0x2C, 0x3E, 0x50)
        matching_body = 0
        total_body = 0

        for p in doc.paragraphs:
            if p.style.name.startswith('Heading'):
                continue
            if not p.text.strip():
                continue
            for run in p.runs:
                if run.text.strip() and len(run.text.strip()) > 5:
                    total_body += 1
                    if run.font.color and run.font.color.rgb:
                        dist = color_distance(run.font.color.rgb, target_body_color)
                        if dist < 30:
                            matching_body += 1
                    break  # check first significant run per paragraph
            if total_body >= 15:
                break  # sample enough

        if total_body > 0 and matching_body / total_body >= 0.5:
            print(f"PASS: Component 9 — Body color #2C3E50 "
                  f"({matching_body}/{total_body} match) (0.05 pts)")
            total_score += 0.05
        elif matching_body > 0:
            print(f"PARTIAL: Component 9 — Some body text matches "
                  f"({matching_body}/{total_body}) (0.02 pts)")
            total_score += 0.02
        else:
            print(f"FAIL: Component 9 — No body text with color #2C3E50 "
                  f"(checked {total_body} paragraphs)")
    except Exception as e:
        print(f"ERROR: Component 9 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    # Try alternate naming
    alt_path = f'{WORKDIR}/Proposal_Complete.docx'
    if os.path.exists(alt_path):
        file_path = alt_path
    else:
        print(f"File not found: {file_path}")
        print("REWARD: 0.0")
        exit()

verify_task(file_path)
