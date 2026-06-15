"""
Reward Script: Vendor Evaluation Report in LibreOffice Writer
Task ID: writer_wf_068
Domain: libreoffice_writer
Scoring:
  Component 1: Title present (0.10)
  Component 2: TOC section (0.10)
  Component 3: Evaluation Summary table (0.20)
  Component 4: Scoring Matrix table (0.20)
  Component 5: Detailed Vendor Assessments - 3 Heading 2 sections (0.15)
  Component 6: Recommendation section (0.10)
  Component 7: Approval Signatures section (0.15)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_068'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts and styles for analysis
    paras = [(p.style.name if p.style else 'Normal', p.text.strip()) for p in doc.paragraphs]
    tables = doc.tables

    # Component 1: Title present with correct text (0.10 points)
    try:
        title_found = False
        for style, text in paras:
            if style == 'Title' and 'vendor evaluation report' in text.lower() and 'it services' in text.lower():
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 — Title 'Vendor Evaluation Report - IT Services' found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Title not found or missing required text")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: TOC section present (0.10 points)
    # The golden doc has a Heading 1 "Table of Contents" — this is task-introduced
    try:
        toc_found = False
        for style, text in paras:
            if 'table of contents' in text.lower():
                toc_found = True
                break
        if toc_found:
            print(f"PASS: Component 2 — TOC section found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — No TOC section found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Evaluation Summary table (0.20 points)
    # Should have a table with headers: Vendor Name, Date, Evaluator
    # At least 5 vendor rows + 1 header row = 6 rows minimum, 3 columns
    try:
        summary_table_found = False
        for table in tables:
            if len(table.rows) >= 2 and len(table.columns) >= 3:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'vendor' in ' '.join(headers) and 'date' in ' '.join(headers) and 'evaluator' in ' '.join(headers):
                    # Check that the table has vendor data rows (at least 3)
                    data_rows = len(table.rows) - 1
                    if data_rows >= 3:
                        summary_table_found = True
                        break
        if summary_table_found:
            print(f"PASS: Component 3 — Evaluation Summary table with vendor data found (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — No valid Evaluation Summary table found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Scoring Matrix table (0.20 points)
    # Should have 5 vendors as rows, 6 criteria columns + vendor name column = 7 cols
    # Criteria: Price, Quality, Delivery, Support, Innovation, Compliance
    # Scores should be 1-5
    try:
        scoring_matrix_found = False
        criteria_keywords = ['price', 'quality', 'delivery', 'support', 'innovation', 'compliance']
        for table in tables:
            if len(table.rows) >= 6 and len(table.columns) >= 7:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                matched_criteria = sum(1 for kw in criteria_keywords if kw in ' '.join(headers))
                if matched_criteria >= 5:
                    # Check that data cells contain numeric scores 1-5
                    valid_scores = 0
                    total_data_cells = 0
                    for ri in range(1, len(table.rows)):
                        for ci in range(1, len(table.columns)):
                            cell_text = table.rows[ri].cells[ci].text.strip()
                            total_data_cells += 1
                            if cell_text.isdigit() and 1 <= int(cell_text) <= 5:
                                valid_scores += 1
                    if total_data_cells > 0 and valid_scores >= total_data_cells * 0.7:
                        scoring_matrix_found = True
                        break
        if scoring_matrix_found:
            print(f"PASS: Component 4 — Scoring Matrix table found with valid scores (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — No valid Scoring Matrix table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Detailed Vendor Assessments with 3 Heading 2 subheadings (0.15 points)
    # Each Heading 2 should be followed by at least one Normal paragraph with content
    try:
        heading2_sections = []
        for i, (style, text) in enumerate(paras):
            if style == 'Heading 2' and text:
                # Check if there's a following paragraph with content
                has_content = False
                for j in range(i + 1, min(i + 5, len(paras))):
                    next_style, next_text = paras[j]
                    if next_style == 'Heading 2' or next_style == 'Heading 1':
                        break
                    if next_text and len(next_text) > 20:
                        has_content = True
                        break
                if has_content:
                    heading2_sections.append(text)

        if len(heading2_sections) >= 3:
            print(f"PASS: Component 5 — Found {len(heading2_sections)} Heading 2 vendor assessment sections (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — Found only {len(heading2_sections)} Heading 2 vendor assessments, need 3")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Recommendation section (0.10 points)
    try:
        recommendation_found = False
        for i, (style, text) in enumerate(paras):
            if 'recommendation' in text.lower() and style in ('Heading 1', 'Heading 2'):
                # Check there's recommendation content after it
                for j in range(i + 1, min(i + 5, len(paras))):
                    _, next_text = paras[j]
                    if next_text and len(next_text) > 20:
                        recommendation_found = True
                        break
                break
        if recommendation_found:
            print(f"PASS: Component 6 — Recommendation section with content found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — No Recommendation section with content found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Approval Signatures section (0.15 points)
    # Should have an "Approval Signatures" heading and signature lines (underscores or "Date:")
    try:
        approval_found = False
        signature_elements = 0
        in_approval_section = False
        for style, text in paras:
            if 'approval' in text.lower() and 'signature' in text.lower():
                in_approval_section = True
                continue
            if in_approval_section:
                if style in ('Heading 1', 'Heading 2') and 'approval' not in text.lower():
                    break  # Left the section
                if '____' in text or 'date:' in text.lower():
                    signature_elements += 1
                elif text and len(text) > 2 and '____' not in text and 'date' not in text.lower():
                    # Name or title line
                    signature_elements += 1

        if in_approval_section and signature_elements >= 4:
            approval_found = True

        if approval_found:
            print(f"PASS: Component 7 — Approval Signatures section found with {signature_elements} elements (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 7 — Approval Signatures section not found or incomplete (in_section={in_approval_section}, elements={signature_elements})")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    final_score = round(final_score, 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
