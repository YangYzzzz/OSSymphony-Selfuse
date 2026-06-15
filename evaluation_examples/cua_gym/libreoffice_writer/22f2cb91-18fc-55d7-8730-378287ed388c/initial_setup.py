"""
Initial Setup: Change tracked changes display colors in LibreOffice Writer
Task ID: writer_rm_015
Domain: libreoffice_writer

Creates a Style_Guide.docx with realistic content and some tracked changes
using default author colors. LibreOffice Writer is opened with the document.
The track changes color settings remain at their defaults (By Author / auto).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Meridian Corp — Brand & Style Guide", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Section 1: Introduction ---
    doc.add_heading("1. Introduction", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "This document establishes the official brand and style guidelines for "
        "Meridian Corporation. All internal communications, marketing materials, "
        "and public-facing content must adhere to these standards to maintain "
        "brand consistency across all channels."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Last reviewed: March 15, 2025. Approved by the Communications Committee."
    )
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Section 2: Brand Voice ---
    doc.add_heading("2. Brand Voice & Tone", level=1)
    doc.add_paragraph(
        "Meridian's brand voice is professional yet approachable. We communicate "
        "with clarity and confidence while remaining empathetic to our audience's "
        "needs. Avoid overly technical jargon unless writing for a specialist audience."
    )

    doc.add_heading("2.1 Key Principles", level=2)
    principles = [
        "Clarity — Write in plain language. Aim for an 8th-grade reading level.",
        "Consistency — Use the approved terminology from the glossary (Section 6).",
        "Confidence — State facts directly. Avoid hedging words like 'maybe' or 'perhaps'.",
        "Empathy — Acknowledge the reader's perspective before presenting solutions.",
    ]
    for principle in principles:
        doc.add_paragraph(principle, style="List Bullet")

    # --- Section 3: Typography ---
    doc.add_heading("3. Typography Standards", level=1)
    doc.add_paragraph(
        "All documents must use the following typeface and sizing conventions:"
    )

    # Table for typography
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    headers = ["Element", "Font", "Size"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    typo_data = [
        ["Headings (H1)", "Calibri Bold", "18pt"],
        ["Headings (H2)", "Calibri Bold", "14pt"],
        ["Body Text", "Calibri", "11pt"],
        ["Footnotes & Captions", "Calibri Light", "9pt"],
    ]
    for r, row_data in enumerate(typo_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Section 4: Color Palette ---
    doc.add_heading("4. Color Palette", level=1)
    doc.add_paragraph(
        "The primary brand colors are used across all marketing and digital assets. "
        "Secondary colors may be used sparingly for accent and emphasis."
    )

    color_table = doc.add_table(rows=5, cols=3)
    color_table.style = "Table Grid"
    color_headers = ["Color Name", "Hex Code", "Usage"]
    for i, h in enumerate(color_headers):
        cell = color_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    color_data = [
        ["Meridian Blue", "#1A5276", "Primary brand color, headers, CTAs"],
        ["Silver Mist", "#BDC3C7", "Backgrounds, subtle dividers"],
        ["Warm Coral", "#E74C3C", "Alerts, emphasis, limited accent use"],
        ["Graphite", "#2C3E50", "Body text, paragraphs"],
    ]
    for r, row_data in enumerate(color_data, 1):
        for c, val in enumerate(row_data):
            color_table.cell(r, c).text = val

    # --- Section 5: Document Formatting ---
    doc.add_heading("5. Document Formatting Rules", level=1)
    doc.add_paragraph(
        "All formal documents follow a standard layout to ensure readability "
        "and a consistent professional appearance."
    )

    formatting_rules = [
        "Page margins: 1 inch on all sides (Letter size, portrait orientation).",
        "Line spacing: 1.15 for body text; single spacing for tables and captions.",
        "Paragraph spacing: 6pt before, 6pt after for body paragraphs.",
        "Page numbers: bottom center, starting from page 2.",
        "Headers: company logo left-aligned; document title right-aligned.",
        "Footnotes: use superscript numerals with Calibri Light 9pt.",
    ]
    for rule in formatting_rules:
        doc.add_paragraph(rule, style="List Number")

    # --- Section 6: Glossary ---
    doc.add_heading("6. Approved Terminology Glossary", level=1)
    doc.add_paragraph(
        "Use the following terms consistently. Unapproved synonyms should be "
        "replaced during the editing process."
    )

    glossary_table = doc.add_table(rows=7, cols=3)
    glossary_table.style = "Table Grid"
    glossary_headers = ["Approved Term", "Do NOT Use", "Context"]
    for i, h in enumerate(glossary_headers):
        cell = glossary_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    glossary_data = [
        ["Client", "Customer / User", "When referring to external partners"],
        ["Team member", "Employee / Worker", "Internal personnel references"],
        ["Solution", "Product / Service", "In marketing materials"],
        ["Onboarding", "Setup / Training", "New client introduction process"],
        ["Deliverable", "Output / Result", "Project milestones and handoffs"],
        ["Stakeholder", "Contact / Person", "Decision-makers in a project"],
    ]
    for r, row_data in enumerate(glossary_data, 1):
        for c, val in enumerate(row_data):
            glossary_table.cell(r, c).text = val

    # --- Section 7: Revision History ---
    doc.add_heading("7. Revision History", level=1)
    doc.add_paragraph(
        "This section tracks changes made to the Style Guide. All modifications "
        "must be reviewed and approved by the Communications Committee before "
        "being incorporated into the official document."
    )

    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        "Note: Track changes should be enabled when editing this document. "
        "Color coding for tracked changes follows the application defaults."
    )
    run_note.font.italic = True
    run_note.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure track changes color settings are at defaults.
    # The default is -1 (By Author) for Insert, Delete, and ChangedAttribute colors.
    # We do NOT modify registrymodifications.xcu for the initial env — defaults apply.

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
