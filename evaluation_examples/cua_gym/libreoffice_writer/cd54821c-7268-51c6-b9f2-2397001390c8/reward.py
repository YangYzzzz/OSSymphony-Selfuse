"""
Reward Script: Delete the paragraph 'TODO: Add screenshots here' from the user guide
Task ID: writer_tech_020
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): 'TODO: Add screenshots here' paragraph is absent
  Component 2 (0.3): Paragraph count is 27 (one fewer than initial 28)
  Component 3 (0.2): Surrounding content flows correctly after deletion
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_020'


def persist_app_state(domain):
    """Attempt to save any unsaved LibreOffice edits via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_texts = [p.text for p in doc.paragraphs]

    # Component 1: 'TODO: Add screenshots here' paragraph is absent (0.5 points)
    # This is the primary task requirement. The placeholder paragraph must be gone.
    try:
        todo_found = any('TODO: Add screenshots here' in t for t in all_texts)
        if not todo_found:
            print(f"PASS: Component 1 -- 'TODO: Add screenshots here' not found in document (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 -- 'TODO: Add screenshots here' still present in document")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Paragraph count is 27 (one fewer than initial 28) (0.3 points)
    # The initial document has 28 paragraphs. After deleting one, should be 27.
    try:
        para_count = len(doc.paragraphs)
        if para_count == 27:
            print(f"PASS: Component 2 -- Paragraph count is 27 as expected (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 -- Expected 27 paragraphs, found {para_count}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Surrounding content flows correctly (0.2 points)
    # After deletion, the paragraph starting with "Run the initialization command"
    # should be followed directly by "After initialization completes".
    # In the initial doc: P20="Run the initialization...", P21="TODO...", P22="After initialization..."
    # In golden: P20="Run the initialization...", P21="After initialization..."
    try:
        # Find the paragraph containing "Run the initialization command"
        run_init_idx = None
        for i, t in enumerate(all_texts):
            if t.startswith("Run the initialization command"):
                run_init_idx = i
                break

        if run_init_idx is not None and run_init_idx + 1 < len(all_texts):
            next_text = all_texts[run_init_idx + 1]
            if next_text.startswith("After initialization completes"):
                print(f"PASS: Component 3 -- Content flows correctly: 'Run the initialization...' -> 'After initialization...' (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 -- After 'Run the initialization...', expected 'After initialization...', found: '{next_text[:60]}'")
        else:
            print(f"FAIL: Component 3 -- Could not find 'Run the initialization command' paragraph")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'

persist_app_state("libreoffice_writer")

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
