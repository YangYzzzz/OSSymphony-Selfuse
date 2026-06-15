"""
Initial Setup: Delete the 'TODO: Add screenshots here' paragraph from the user guide.
Task ID: writer_tech_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_020'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('StreamDB Installation and Configuration Guide', level=0)

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'StreamDB is a high-performance distributed database designed for real-time '
        'analytics workloads. This guide walks you through the complete installation '
        'process, from system requirements to cluster configuration.'
    )
    doc.add_paragraph(
        'StreamDB supports horizontal scaling across multiple nodes and provides '
        'automatic failover with configurable replication factors. It is optimized '
        'for time-series data and supports SQL-compatible query syntax.'
    )

    # --- Section 2: System Requirements ---
    doc.add_heading('2. System Requirements', level=1)
    doc.add_paragraph('Before installing StreamDB, ensure your system meets the following minimum requirements:')

    # Requirements as bullet list
    doc.add_paragraph('Operating System: Ubuntu 22.04 LTS, RHEL 8+, or Debian 12+', style='List Bullet')
    doc.add_paragraph('CPU: 4 cores minimum (8 cores recommended for production)', style='List Bullet')
    doc.add_paragraph('RAM: 16 GB minimum (32 GB recommended)', style='List Bullet')
    doc.add_paragraph('Disk: 100 GB SSD with at least 500 IOPS', style='List Bullet')
    doc.add_paragraph('Network: 1 Gbps dedicated interface for cluster communication', style='List Bullet')
    doc.add_paragraph('Java Runtime: OpenJDK 17 or later', style='List Bullet')

    # --- Section 3: Installation Steps ---
    doc.add_heading('3. Installation Steps', level=1)

    doc.add_heading('3.1 Download and Extract', level=2)
    doc.add_paragraph(
        'Download the latest StreamDB release from the official repository. '
        'The package includes the server binary, CLI tools, and default configuration files.'
    )
    # Code-like paragraph
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(
        'wget https://releases.streamdb.io/v3.2.1/streamdb-3.2.1-linux-amd64.tar.gz\n'
        'tar xzf streamdb-3.2.1-linux-amd64.tar.gz\n'
        'sudo mv streamdb-3.2.1 /opt/streamdb'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_heading('3.2 Configure Environment Variables', level=2)
    doc.add_paragraph(
        'Add StreamDB to your system PATH and set the data directory. '
        'Edit your shell profile (~/.bashrc or ~/.zshrc) and append the following lines:'
    )
    env_para = doc.add_paragraph()
    env_run = env_para.add_run(
        'export STREAMDB_HOME=/opt/streamdb\n'
        'export PATH=$PATH:$STREAMDB_HOME/bin\n'
        'export STREAMDB_DATA=/var/lib/streamdb/data'
    )
    env_run.font.name = 'Courier New'
    env_run.font.size = Pt(9)

    doc.add_heading('3.3 Initialize the Database', level=2)
    doc.add_paragraph(
        'Run the initialization command to set up the data directory structure '
        'and generate the default configuration. This step creates the WAL journal, '
        'metadata catalog, and default keyspace.'
    )

    # *** THE TODO PARAGRAPH (to be deleted by agent) ***
    todo_para = doc.add_paragraph('TODO: Add screenshots here')

    doc.add_paragraph(
        'After initialization completes, verify the directory structure was created correctly '
        'by listing the contents of the data directory. You should see subdirectories for '
        'wal/, meta/, and data/.'
    )

    # --- Section 4: Cluster Configuration ---
    doc.add_heading('4. Cluster Configuration', level=1)
    doc.add_paragraph(
        'For production deployments, StreamDB should be configured as a multi-node cluster '
        'with a minimum of three nodes for fault tolerance. Each node requires a unique '
        'node ID and must be able to communicate with all other nodes on the designated '
        'cluster port (default: 9420).'
    )

    # Configuration table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Parameter', 'Default Value', 'Description']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    config_data = [
        ['cluster.name', 'streamdb-cluster', 'Unique name for the cluster'],
        ['node.id', 'auto', 'Unique identifier for this node'],
        ['replication.factor', '3', 'Number of data replicas across nodes'],
        ['cluster.port', '9420', 'Port for inter-node communication'],
        ['max.connections', '256', 'Maximum concurrent client connections'],
    ]
    for r, row_data in enumerate(config_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Section 5: Monitoring ---
    doc.add_heading('5. Monitoring and Health Checks', level=1)
    doc.add_paragraph(
        'StreamDB exposes a built-in metrics endpoint on port 9421 that is compatible '
        'with Prometheus scrapers. Key metrics to monitor include query latency percentiles, '
        'replication lag, disk utilization, and connection pool saturation.'
    )
    doc.add_paragraph(
        'For alerting, we recommend setting thresholds on p99 query latency (alert if > 200ms), '
        'replication lag (alert if > 5 seconds), and disk usage (alert if > 80%). '
        'Integration with Grafana dashboards is available via the official StreamDB plugin.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
