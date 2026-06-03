"""
Reward Script: Insert a row between Phase 2 and Phase 3 with light purple background
Task ID: writer_tbl_063
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table has 6 rows after inserting new row between Phase 2 and Phase 3
  Component 2 (0.4): New row 4 contains correct content: Phase 2.5, Integration Testing, 2024-06-01, 2024-06-15
  Component 3 (0.3): New row has light purple background (fill color ~DDA0DD or similar purple)
"""

import os
from math import sqrt

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_063'


def color_distance(hex1, hex2):
    """Compute Euclidean RGB distance between two hex color strings."""
    r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
    r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def get_cell_fill_color(cell):
    """Return the fill hex color string for a table cell, or None if no shading."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    # 'auto', None, or '000000' with val=clear but fill present
    if fill and fill.upper() not in ('AUTO', '000000'):
        return fill.upper()
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Insert a row between 'Phase 2' and 'Phase 3'. Fill it with
    'Phase 2.5', 'Integration Testing', '2024-06-01', '2024-06-15'.
    Then set the background of this new row to light purple.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least 1 table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)

    # Component 1: Table has 6 rows (0.3 points)
    # Initial has 5 rows; golden has 6 rows after inserting new row
    try:
        if num_rows == 6:
            print(f"PASS: Component 1 — Table has 6 rows as expected (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — Expected 6 rows, found {num_rows}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row 4 (index 3) contains correct content (0.4 points)
    # Expected: Phase 2.5 | Integration Testing | 2024-06-01 | 2024-06-15
    expected_cells = ['Phase 2.5', 'Integration Testing', '2024-06-01', '2024-06-15']
    try:
        if num_rows >= 4:
            new_row = table.rows[3]
            actual_cells = [cell.text.strip() for cell in new_row.cells]
            if actual_cells == expected_cells:
                print(f"PASS: Component 2 — Row 4 content matches: {actual_cells} (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — Expected {expected_cells}, found {actual_cells}")
        else:
            print(f"FAIL: Component 2 — Not enough rows to check row 4")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: New row (row index 3) has a light purple background (0.3 points)
    # Light purple reference: DDA0DD (plum), or nearby purple shades
    # Accept any color within Euclidean RGB distance of ~80 from DDA0DD (plum)
    # or from CC99FF / B39DDB or other standard light purples
    LIGHT_PURPLE_REFS = ['DDA0DD', 'CC99FF', 'B39DDB', 'D8B4FE', 'C9B1FF', 'E6CCFF', 'D9B3FF', 'CBA5E8']
    PURPLE_DISTANCE_THRESHOLD = 80
    try:
        if num_rows >= 4:
            new_row = table.rows[3]
            row_fills = []
            for cell in new_row.cells:
                fill = get_cell_fill_color(cell)
                row_fills.append(fill)

            # Check if all cells have a fill color
            non_none_fills = [f for f in row_fills if f is not None]
            if not non_none_fills:
                print(f"FAIL: Component 3 — New row has no background fill color (found {row_fills})")
            else:
                # Check if the fill color is close to a light purple reference
                first_fill = non_none_fills[0]
                is_purple = any(
                    color_distance(first_fill, ref) <= PURPLE_DISTANCE_THRESHOLD
                    for ref in LIGHT_PURPLE_REFS
                )
                # Also check: it should be in the purple/violet range: R>B>G roughly, or R~B and low G
                r = int(first_fill[0:2], 16)
                g = int(first_fill[2:4], 16)
                b = int(first_fill[4:6], 16)
                is_purple_heuristic = (r > g) and (b > g) and (r >= 150) and (b >= 150)

                if is_purple or is_purple_heuristic:
                    print(f"PASS: Component 3 — New row has light purple background fill={first_fill} (0.3 pts)")
                    total_score += 0.3
                else:
                    print(f"FAIL: Component 3 — New row background fill={first_fill} is not light purple")
        else:
            print(f"FAIL: Component 3 — Not enough rows to check row 4 background")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
