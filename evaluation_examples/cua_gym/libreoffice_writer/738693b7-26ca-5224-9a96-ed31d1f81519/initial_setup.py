"""
Initial Setup: Project phases table - insert row task
Task ID: writer_tbl_063
Domain: libreoffice_writer

Creates a .docx with a 5-row x 4-column project phases table.
The table does NOT yet have the Phase 2.5 row or any purple background.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP_PATH = f'{WORKDIR}/Desktop/project_phases.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title paragraph for context
    title = doc.add_paragraph()
    run = title.add_run('Project Phases Overview')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph('')  # blank line

    # Create the 5-row x 4-column table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Define table data: header + 4 phases (no Phase 2.5)
    rows_data = [
        ('Phase',   'Activity',         'Start',      'End'),
        ('Phase 1', 'Requirements',     '2024-01-15', '2024-02-28'),
        ('Phase 2', 'Development',      '2024-03-01', '2024-05-31'),
        ('Phase 3', 'Deployment',       '2024-07-01', '2024-07-31'),
        ('Phase 4', 'Maintenance',      '2024-08-01', '2024-12-31'),
    ]

    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            # Clear default paragraph and set text
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            if r_idx == 0:
                run.bold = True  # Header row bold

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also place the file on Desktop as project_phases.docx (per task context)
    import shutil
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)
    shutil.copy(OUTPUT, DESKTOP_PATH)
    print(f'Copied to Desktop: {DESKTOP_PATH}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
