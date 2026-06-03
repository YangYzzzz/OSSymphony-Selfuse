"""
Initial Setup: Theater program cast list with inconsistent spacing
Task ID: writer_rd_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Riverside Community Theater', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Presents')
    run.font.size = Pt(14)
    run.font.italic = True

    play_title = doc.add_paragraph()
    play_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = play_title.add_run('The Merchant of Venice')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    season = doc.add_paragraph()
    season.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = season.add_run('Spring Season 2025')
    run.font.size = Pt(12)
    run.font.italic = True

    doc.add_paragraph()  # blank line

    # --- Director's Note ---
    doc.add_heading("Director's Note", level=1)
    note = doc.add_paragraph(
        "This production reimagines Shakespeare's complex exploration of justice, "
        "mercy, and identity in a contemporary setting. Our talented ensemble brings "
        "fresh perspectives to these timeless characters while honoring the beauty of "
        "the original text. We hope you enjoy this evening's performance."
    )
    note.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()  # blank line

    # --- Cast Section ---
    doc.add_heading('Cast', level=1)

    # 8 cast lines with INCONSISTENT spacing (no tab stops)
    # This is the "before" state - irregular spaces make alignment messy
    cast_lines = [
        ('Antonio', 'Richard Blackwell'),
        ('Portia', 'Elena Vasquez'),
        ('Shylock', 'David Chen'),
        ('Bassanio', 'Thomas Grant'),
        ('Nerissa', 'Amelia Foster'),
        ('Gratiano', 'Marcus Okonkwo'),
        ('Jessica', 'Sarah Lindqvist'),
        ('Lorenzo', 'James Whitfield'),
    ]

    for character, actor in cast_lines:
        para = doc.add_paragraph()
        # Use irregular spacing to simulate a poorly formatted cast list
        spacing = ' ' * (24 - len(character))
        run = para.add_run(f'{character}{spacing}{actor}')
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()  # blank line

    # --- Production Team ---
    doc.add_heading('Production Team', level=1)

    team = [
        ('Director', 'Patricia Hawthorne'),
        ('Stage Manager', 'Kevin Liu'),
        ('Set Designer', 'Maria Gonzalez'),
        ('Costume Designer', 'Yuki Tanaka'),
        ('Lighting Designer', 'Robert Singh'),
        ('Sound Designer', 'Anna Kowalski'),
    ]

    for role, name in team:
        para = doc.add_paragraph()
        run_role = para.add_run(f'{role}: ')
        run_role.font.bold = True
        run_role.font.size = Pt(11)
        run_name = para.add_run(name)
        run_name.font.size = Pt(11)

    doc.add_paragraph()  # blank line

    # --- Venue Information ---
    doc.add_heading('Venue Information', level=1)
    venue_info = doc.add_paragraph(
        'Riverside Community Theater\n'
        '245 Elm Street, Riverside, CA 92501\n'
        'Box Office: (951) 555-0142\n'
        'www.riversidetheater.org'
    )
    venue_info.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
