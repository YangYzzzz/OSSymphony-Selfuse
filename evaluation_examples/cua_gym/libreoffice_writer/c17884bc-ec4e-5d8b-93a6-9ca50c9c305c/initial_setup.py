"""
Initial Setup: Project Timeline Document (No table yet)
Task ID: writer_tbl_080
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_080'
OUTPUT = f'{WORKDIR}/Desktop/project_timeline.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # Title: "Project Alpha - Timeline" — bold, 16pt, centered
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("Project Alpha - Timeline")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Subtitle line
    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run("Version 1.0 | Last Updated: March 2024")
    subtitle_run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
