"""
Reward Script: Build a complete project timeline table in a .docx file
Task ID: writer_tbl_080
Domain: libreoffice_writer
Scoring:
  Component 1: Table exists with 8 rows x 5 columns and correct content (0.30 pts)
  Component 2: Header row has bold white text + dark blue (#1F3664) background (0.30 pts)
  Component 3: Vertical cell merges in column A (Planning/Execution/Testing) (0.20 pts)
  Component 4: Outer border dark blue 2pt + inner borders gray 0.5pt + table centered (0.20 pts)
Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_080'

FILE_PATH = '/home/user/Desktop/project_timeline.docx'

# Expected header columns
EXPECTED_HEADERS = ['Phase', 'Task', 'Start Date', 'End Date', 'Status']

# Expected data in subsequent rows (col1..4; col0 is the merged Phase)
EXPECTED_ROWS = [
    # (task, start, end, status)
    ('Requirements Gathering', '2024-01-15', '2024-02-15', 'Complete'),
    ('Design Review',          '2024-02-16', '2024-03-15', 'Complete'),
    ('Backend Development',    '2024-03-16', '2024-05-15', 'In Progress'),
    ('Frontend Development',   '2024-04-01', '2024-05-31', 'In Progress'),
    ('Unit Testing',           '2024-05-01', '2024-06-15', 'Not Started'),
    ('Integration Testing',    '2024-06-01', '2024-06-30', 'Not Started'),
    ('Go Live',                '2024-07-01', '2024-07-15', 'Not Started'),
]


def color_close(actual_hex, expected_hex, tolerance=30):
    """Check if two hex colors are perceptually close."""
    try:
        ar, ag, ab = int(actual_hex[0:2], 16), int(actual_hex[2:4], 16), int(actual_hex[4:6], 16)
        er, eg, eb = int(expected_hex[0:2], 16), int(expected_hex[2:4], 16), int(expected_hex[4:6], 16)
        dist = ((ar - er)**2 + (ag - eg)**2 + (ab - eb)**2) ** 0.5
        return dist <= tolerance
    except Exception:
        return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Table exists with 8 rows x 5 columns and correct content
    # (0.30 points)
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 — no table found in document")
        else:
            t = doc.tables[0]
            nrows = len(t.rows)
            ncols = len(t.columns)

            if nrows != 8 or ncols != 5:
                print(f"FAIL: Component 1 — expected 8x5 table, found {nrows}x{ncols}")
            else:
                # Verify headers
                header_row = [t.cell(0, c).text.strip() for c in range(5)]
                headers_ok = (header_row == EXPECTED_HEADERS)

                # Verify data rows (rows 1-7)
                data_ok = True
                mismatches = []
                for ri, (exp_task, exp_start, exp_end, exp_status) in enumerate(EXPECTED_ROWS):
                    row_idx = ri + 1
                    actual_task   = t.cell(row_idx, 1).text.strip()
                    actual_start  = t.cell(row_idx, 2).text.strip()
                    actual_end    = t.cell(row_idx, 3).text.strip()
                    actual_status = t.cell(row_idx, 4).text.strip()
                    if actual_task != exp_task or actual_start != exp_start \
                       or actual_end != exp_end or actual_status != exp_status:
                        data_ok = False
                        mismatches.append(
                            f"Row{row_idx}: expected ({exp_task},{exp_start},{exp_end},{exp_status})"
                            f" got ({actual_task},{actual_start},{actual_end},{actual_status})"
                        )

                if headers_ok and data_ok:
                    print("PASS: Component 1 — 8x5 table with correct headers and all 7 data rows (0.30 pts)")
                    total_score += 0.30
                elif headers_ok and not data_ok:
                    # Partial: headers correct but some data rows wrong
                    print(f"PARTIAL: Component 1 — headers correct but data mismatches: {mismatches[:3]}")
                    # Award half credit for correct structure + headers
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 1 — headers wrong: expected {EXPECTED_HEADERS}, got {header_row}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Header row bold white text + dark blue (#1F3664) background
    # (0.30 points)
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — no table found")
        else:
            t = doc.tables[0]
            if len(t.rows) < 1:
                print("FAIL: Component 2 — table has no rows")
            else:
                bold_white_count = 0
                blue_bg_count = 0

                for ci in range(min(5, len(t.columns))):
                    cell = t.cell(0, ci)
                    # Check background fill via XML
                    tcp = cell._tc.find(qn('w:tcPr'))
                    if tcp is not None:
                        shd = tcp.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'), '').upper()
                            if color_close(fill, '1F3664', tolerance=40):
                                blue_bg_count += 1
                    # Check bold + white text
                    for para in cell.paragraphs:
                        for run in para.runs:
                            is_bold = run.font.bold is True
                            try:
                                rgb = run.font.color.rgb
                                color_hex = str(rgb).upper() if rgb else None
                            except Exception:
                                color_hex = None
                            if is_bold and color_hex and color_close(color_hex, 'FFFFFF', tolerance=10):
                                bold_white_count += 1

                if blue_bg_count >= 5 and bold_white_count >= 5:
                    print(f"PASS: Component 2 — all 5 header cells have dark blue bg and bold white text (0.30 pts)")
                    total_score += 0.30
                elif blue_bg_count >= 3 or bold_white_count >= 3:
                    print(f"PARTIAL: Component 2 — blue_bg_cells={blue_bg_count}/5, bold_white_runs={bold_white_count}/5")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 2 — blue_bg_cells={blue_bg_count}/5, bold_white_runs={bold_white_count}/5")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Vertical merges in column A: A2-A3=Planning, A4-A5=Execution,
    #              A6-A7=Testing, A8=Launch (0.20 points)
    #
    # Merge detection via raw XML vMerge:
    #   - 'restart' = first cell of a vertical merge group
    #   - None (present but no val) = continuation cell (merged into above)
    #   - absent = standalone cell (no merge)
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — no table found")
        else:
            t = doc.tables[0]
            if len(t.rows) < 8:
                print("FAIL: Component 3 — table has only %d rows, need 8" % len(t.rows))
            else:
                def get_vmerge(table, row_idx):
                    """Return vMerge value for column-0 cell at row_idx.
                    Returns: 'restart', 'continue', or 'none' (no vMerge element)."""
                    tr = table.rows[row_idx]._tr
                    tcs = tr.findall(qn('w:tc'))
                    if not tcs:
                        return 'none'
                    tc = tcs[0]
                    tcp = tc.find(qn('w:tcPr'))
                    if tcp is None:
                        return 'none'
                    vm = tcp.find(qn('w:vMerge'))
                    if vm is None:
                        return 'none'
                    val = vm.get(qn('w:val'))
                    return 'restart' if val == 'restart' else 'continue'

                def get_col0_text(table, row_idx):
                    """Get text of the first tc in a row from raw XML."""
                    tr = table.rows[row_idx]._tr
                    tcs = tr.findall(qn('w:tc'))
                    if not tcs:
                        return ''
                    wts = tcs[0].findall('.//' + qn('w:t'))
                    return ''.join(wt.text or '' for wt in wts).strip()

                # Check merge pattern: rows 1-2 Planning, 3-4 Execution, 5-6 Testing, 7 Launch
                planning_merged  = (get_vmerge(t, 1) == 'restart' and get_vmerge(t, 2) == 'continue')
                execution_merged = (get_vmerge(t, 3) == 'restart' and get_vmerge(t, 4) == 'continue')
                testing_merged   = (get_vmerge(t, 5) == 'restart' and get_vmerge(t, 6) == 'continue')
                launch_separate  = (get_vmerge(t, 7) == 'none')

                # Check text in the first cell of each merge group
                planning_text   = get_col0_text(t, 1)
                execution_text  = get_col0_text(t, 3)
                testing_text    = get_col0_text(t, 5)
                launch_text     = get_col0_text(t, 7)

                planning_text_ok  = (planning_text == 'Planning')
                execution_text_ok = (execution_text == 'Execution')
                testing_text_ok   = (testing_text == 'Testing')
                launch_text_ok    = (launch_text == 'Launch')

                merge_score = sum([planning_merged, execution_merged, testing_merged, launch_separate])
                text_score  = sum([planning_text_ok, execution_text_ok, testing_text_ok, launch_text_ok])

                if merge_score == 4 and text_score == 4:
                    print("PASS: Component 3 — all 4 phase cells correctly merged/unmerged with correct text (0.20 pts)")
                    total_score += 0.20
                elif merge_score >= 3 and text_score >= 3:
                    print("PARTIAL: Component 3 — merges=%d/4, texts=%d/4" % (merge_score, text_score))
                    total_score += 0.10
                else:
                    print("FAIL: Component 3 — merges=%d/4 "
                          "(planning=%s vmerge=%s, exec=%s vmerge=%s, testing=%s vmerge=%s), "
                          "texts=%d/4 "
                          "(planning=%s=%r, exec=%s=%r, testing=%s=%r, launch=%s=%r)" % (
                              merge_score,
                              planning_merged,  '%s/%s' % (get_vmerge(t,1), get_vmerge(t,2)),
                              execution_merged, '%s/%s' % (get_vmerge(t,3), get_vmerge(t,4)),
                              testing_merged,   '%s/%s' % (get_vmerge(t,5), get_vmerge(t,6)),
                              text_score,
                              planning_text_ok,  planning_text,
                              execution_text_ok, execution_text,
                              testing_text_ok,   testing_text,
                              launch_text_ok,    launch_text,
                          ))
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Table borders (outer 2pt dark blue, inner 0.5pt gray)
    #              AND table centered on page (0.20 points)
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 4 — no table found")
        else:
            t = doc.tables[0]
            tblPr = t._tbl.find(qn('w:tblPr'))

            # Check table alignment (centering)
            table_centered = False
            if tblPr is not None:
                jc = tblPr.find(qn('w:jc'))
                if jc is not None:
                    table_centered = (jc.get(qn('w:val')) == 'center')

            # Check borders
            outer_border_ok = False
            inner_border_ok = False

            if tblPr is not None:
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is not None:
                    # Outer borders: top, left, bottom, right — sz=16 (~2pt), color close to 1F3664
                    outer_sides = ['top', 'left', 'bottom', 'right']
                    outer_ok_count = 0
                    for side in outer_sides:
                        el = tblBorders.find(qn(f'w:{side}'))
                        if el is not None:
                            sz = el.get(qn('w:sz'), '0')
                            color = el.get(qn('w:color'), '').upper()
                            # sz=16 means 2pt (in eighths of a point: 16/8=2)
                            sz_ok = (int(sz) >= 14)  # allow slight variation
                            color_ok = color_close(color, '1F3664', tolerance=40)
                            if sz_ok and color_ok:
                                outer_ok_count += 1
                    outer_border_ok = (outer_ok_count >= 3)

                    # Inner borders: insideH, insideV — sz=4 (~0.5pt), color close to gray A0A0A0
                    inner_sides = ['insideH', 'insideV']
                    inner_ok_count = 0
                    for side in inner_sides:
                        el = tblBorders.find(qn(f'w:{side}'))
                        if el is not None:
                            sz = el.get(qn('w:sz'), '0')
                            color = el.get(qn('w:color'), '').upper()
                            # sz=4 means 0.5pt
                            sz_ok = (int(sz) <= 6)  # small inner border
                            color_ok = color_close(color, 'A0A0A0', tolerance=60)
                            if sz_ok and color_ok:
                                inner_ok_count += 1
                    inner_border_ok = (inner_ok_count >= 1)

            checks = [table_centered, outer_border_ok, inner_border_ok]
            passed = sum(checks)

            if passed == 3:
                print("PASS: Component 4 — table centered, outer 2pt dark blue borders, inner 0.5pt gray borders (0.20 pts)")
                total_score += 0.20
            elif passed >= 2:
                print(f"PARTIAL: Component 4 — centered={table_centered}, outer_border={outer_border_ok}, inner_border={inner_border_ok}")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — centered={table_centered}, outer_border={outer_border_ok}, inner_border={inner_border_ok}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
