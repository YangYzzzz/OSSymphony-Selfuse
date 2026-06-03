"""
Reward Script: Landscape Table page style with page break
Task ID: writer_bs_087
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.15): Document has multiple sections (>= 3)
  - Component 2 (0.20): A landscape section exists with A4 dimensions
  - Component 3 (0.20): Landscape section has ~1.5cm margins (all four)
  - Component 4 (0.10): Landscape section header is disabled (empty, not linked)
  - Component 5 (0.20): Landscape section footer has centered PAGE field
  - Component 6 (0.15): Section break inserted before the table paragraph
"""

import os
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_087'

# Tolerance for margins: 1.5cm = 540000 EMU, allow +/- 10000
MARGIN_TARGET = 540000
MARGIN_TOL = 15000

# A4 landscape: w ~= 16838 twips (10692130 EMU), h ~= 11906 twips (7560310 EMU)
# A4 portrait:  w ~= 11906 twips (7560310 EMU),  h ~= 16838 twips (10692130 EMU)
A4_LONG = 10692000   # ~29.7cm in EMU
A4_SHORT = 7560000   # ~21cm in EMU
DIM_TOL = 50000      # tolerance for dimension check


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = doc.sections
    num_sections = len(sections)

    # Component 1: Document has multiple sections (>= 3) — 0.15 points
    # Initial doc has only 1 section; golden has 3 (portrait, landscape, portrait)
    try:
        if num_sections >= 3:
            print(f"PASS: Component 1 — Document has {num_sections} sections (>= 3) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Document has {num_sections} sections, expected >= 3")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Find the landscape section (should be section index 1 or later)
    landscape_idx = None
    for i, s in enumerate(sections):
        if s.orientation == WD_ORIENT.LANDSCAPE:
            landscape_idx = i
            break

    # Component 2: A landscape section exists with A4 dimensions — 0.20 points
    try:
        if landscape_idx is not None:
            ls = sections[landscape_idx]
            w = ls.page_width
            h = ls.page_height
            # Landscape A4: width ~ A4_LONG, height ~ A4_SHORT
            w_ok = abs(w - A4_LONG) < DIM_TOL
            h_ok = abs(h - A4_SHORT) < DIM_TOL
            if w_ok and h_ok:
                print(f"PASS: Component 2 — Landscape section {landscape_idx} with A4 dims (w={w}, h={h}) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — Landscape section {landscape_idx} dims wrong: w={w} (exp ~{A4_LONG}), h={h} (exp ~{A4_SHORT})")
        else:
            print(f"FAIL: Component 2 — No landscape section found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Landscape section margins are ~1.5cm — 0.20 points
    try:
        if landscape_idx is not None:
            ls = sections[landscape_idx]
            margins = {
                'left': ls.left_margin,
                'right': ls.right_margin,
                'top': ls.top_margin,
                'bottom': ls.bottom_margin,
            }
            all_ok = all(abs(v - MARGIN_TARGET) < MARGIN_TOL for v in margins.values())
            if all_ok:
                print(f"PASS: Component 3 — All margins ~1.5cm: {margins} (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Margins not ~1.5cm: {margins} (target={MARGIN_TARGET}, tol={MARGIN_TOL})")
        else:
            print(f"FAIL: Component 3 — No landscape section to check margins")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Landscape section header is disabled — 0.10 points
    # Header should be empty/not linked (header disabled means no content)
    try:
        if landscape_idx is not None:
            ls = sections[landscape_idx]
            header = ls.header
            # Header is "disabled" = either linked_to_previous=False with empty content,
            # or no real header content. Key: no visible header text.
            header_text = ""
            if header.paragraphs:
                header_text = "".join(p.text for p in header.paragraphs).strip()
            header_empty = (header_text == "")
            # Also check it's not linked to previous (independent section header)
            not_linked = not header.is_linked_to_previous
            if header_empty and not_linked:
                print(f"PASS: Component 4 — Header disabled (empty, not linked to prev) (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Header not disabled: text={header_text!r}, linked={header.is_linked_to_previous}")
        else:
            print(f"FAIL: Component 4 — No landscape section to check header")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Landscape section footer has centered PAGE field — 0.20 points
    try:
        if landscape_idx is not None:
            ls = sections[landscape_idx]
            footer = ls.footer
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Check footer is not linked to previous
            footer_independent = not footer.is_linked_to_previous

            # Check for PAGE field in footer
            has_page_field = any(
                'PAGE' in (instr.text or '')
                for p in footer.paragraphs
                for instr in p._element.findall('.//w:instrText', ns)
            )

            # Check centering via alignment API or XML jc element
            is_centered = any(
                p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                or (p._element.find('.//w:jc', ns) is not None
                    and p._element.find('.//w:jc', ns).get(qn('w:val')) == 'center')
                for p in footer.paragraphs
            )

            if footer_independent and has_page_field and is_centered:
                print(f"PASS: Component 5 — Footer has centered PAGE field (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 5 — Footer issues: independent={footer_independent}, page_field={has_page_field}, centered={is_centered}")
        else:
            print(f"FAIL: Component 5 — No landscape section to check footer")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Section break before table paragraph — 0.15 points
    # In the golden doc, paragraph before the table has a sectPr in its pPr (section break)
    try:
        # Find paragraphs that have sectPr in their pPr
        paras_with_sectpr = []
        for i, p in enumerate(doc.paragraphs):
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                sectPr_list = pPr.findall(qn('w:sectPr'))
                if sectPr_list:
                    paras_with_sectpr.append(i)

        # We expect at least one sectPr in a paragraph before or near the table
        # In golden, P3 and P4 have sectPr. The key check: at least one paragraph
        # has a section break, and the document has the landscape section in the middle.
        if len(paras_with_sectpr) >= 1 and landscape_idx is not None:
            print(f"PASS: Component 6 — Section break(s) found at paragraphs {paras_with_sectpr} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 6 — No section break paragraphs found (sectPr in pPr): {paras_with_sectpr}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
