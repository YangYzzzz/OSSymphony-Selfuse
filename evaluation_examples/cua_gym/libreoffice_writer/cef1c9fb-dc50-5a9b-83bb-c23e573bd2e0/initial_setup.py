"""
Initial Setup: Create a Writer document with portrait orientation and a wide table
Task ID: writer_bs_087
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_087'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default portrait, A4, standard margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Title
    heading = doc.add_heading("Quarterly Regional Sales Report — FY2025 Q3", level=1)

    # Introductory paragraphs
    doc.add_paragraph(
        "This report summarizes the sales performance across all regional offices "
        "for the third quarter of fiscal year 2025. The data covers July through "
        "September and includes breakdowns by product category, representative, "
        "and territory."
    )
    doc.add_paragraph(
        "Overall revenue grew 12.4% compared to Q2, driven primarily by strong "
        "performance in the Western and Southern regions. The Eastern region "
        "experienced a slight decline due to delayed contract renewals with two "
        "major enterprise clients."
    )
    doc.add_paragraph(
        "The following table provides detailed figures for each sales representative "
        "across all regions. Note that the table is wide due to the number of "
        "monthly columns and category breakdowns."
    )

    # Wide table with lots of columns — this is the table that needs landscape
    headers = [
        "Rep Name", "Region", "Territory",
        "Jul Hardware", "Jul Software", "Jul Services",
        "Aug Hardware", "Aug Software", "Aug Services",
        "Sep Hardware", "Sep Software", "Sep Services",
        "Q3 Total"
    ]

    data_rows = [
        ["Sarah Chen", "Western", "CA-North",
         "45230", "32100", "18750",
         "48900", "35600", "19200",
         "52100", "38400", "21500",
         "311780"],
        ["Marcus Johnson", "Western", "CA-South",
         "38400", "28900", "15600",
         "41200", "31500", "16800",
         "43800", "33200", "17900",
         "267300"],
        ["Priya Patel", "Eastern", "NY-Metro",
         "52300", "41200", "22100",
         "49800", "38600", "20500",
         "47200", "36100", "19800",
         "327600"],
        ["David Kim", "Eastern", "NJ-Central",
         "31200", "24500", "13800",
         "29800", "22900", "12600",
         "28500", "21700", "12100",
         "197100"],
        ["Elena Rodriguez", "Southern", "TX-Dallas",
         "42800", "33500", "19200",
         "46100", "36200", "20800",
         "49500", "38900", "22400",
         "309400"],
        ["James Walker", "Southern", "FL-Miami",
         "36500", "27800", "16100",
         "39200", "30100", "17500",
         "41800", "32500", "18700",
         "260200"],
        ["Lisa Thompson", "Central", "IL-Chicago",
         "44100", "34200", "18900",
         "47300", "36800", "20300",
         "50200", "39100", "21700",
         "312600"],
        ["Robert Martinez", "Central", "OH-Columbus",
         "29800", "22100", "12500",
         "31500", "23800", "13200",
         "33200", "25100", "14000",
         "205200"],
        ["Amanda Foster", "Western", "WA-Seattle",
         "41600", "31800", "17400",
         "44200", "34100", "18600",
         "47100", "36500", "19800",
         "291100"],
        ["Thomas Wright", "Eastern", "MA-Boston",
         "37900", "29200", "16300",
         "35600", "27100", "15200",
         "34100", "25800", "14500",
         "235700"],
        ["Michelle Lee", "Southern", "GA-Atlanta",
         "33800", "26100", "14700",
         "36200", "28400", "15900",
         "38700", "30200", "17100",
         "241100"],
        ["Kevin Brown", "Central", "MN-Minneapolis",
         "27500", "20800", "11600",
         "29100", "22200", "12400",
         "30800", "23500", "13200",
         "191100"],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            # Right-align numeric columns
            if i >= 3:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Post-table commentary
    doc.add_paragraph("")
    doc.add_paragraph(
        "Key Observations: The Western region continues to lead in total revenue, "
        "with Sarah Chen and Amanda Foster both exceeding their quarterly targets. "
        "The Central region showed consistent month-over-month growth across all "
        "product categories."
    )
    doc.add_paragraph(
        "Action Items: Schedule regional review meetings for October. Prepare "
        "contract renewal proposals for Eastern region enterprise accounts. "
        "Evaluate Southern region expansion opportunities in Nashville and Charlotte markets."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
