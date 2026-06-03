"""
Reward Script: Format compensation section as a structured table
Task ID: writer_hr_052
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Table exists with correct dimensions (2 cols, 7 rows)
  Component 2 (0.30): Table contains correct compensation data
  Component 3 (0.20): Header row has gray shading (dark fill)
  Component 4 (0.20): Alternating row shading on data rows
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_052'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_cell_fill(cell):
    """Extract the w:fill attribute from cell shading, or None."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    return shd.get(qn('w:fill'))


def is_dark_gray_fill(fill_hex):
    """Check if a hex fill color is a dark gray suitable for a header row."""
    if fill_hex is None:
        return False
    try:
        val = int(fill_hex, 16)
        r = (val >> 16) & 0xFF
        g = (val >> 8) & 0xFF
        b = val & 0xFF
        # Dark gray: all channels roughly equal and relatively low (< 128)
        avg = (r + g + b) / 3
        spread = max(r, g, b) - min(r, g, b)
        return avg < 140 and spread < 60
    except (ValueError, TypeError):
        return False


def is_light_gray_fill(fill_hex):
    """Check if a hex fill color is a light gray suitable for alternating shading."""
    if fill_hex is None:
        return False
    try:
        val = int(fill_hex, 16)
        r = (val >> 16) & 0xFF
        g = (val >> 8) & 0xFF
        b = val & 0xFF
        # Light gray: all channels roughly equal and high (> 180)
        avg = (r + g + b) / 3
        spread = max(r, g, b) - min(r, g, b)
        return avg > 180 and spread < 40
    except (ValueError, TypeError):
        return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document — task requires a compensation table")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table has correct dimensions — 2 columns and 7 rows (0.30 points)
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_cols == 2 and num_rows == 7:
            print(f"PASS: Component 1 — Table dimensions correct: {num_rows} rows x {num_cols} cols (0.30 pts)")
            total_score += 0.30
        elif num_cols == 2 and num_rows >= 6:
            # Partial credit: right columns, close row count
            print(f"PARTIAL: Component 1 — Table has 2 cols but {num_rows} rows (expected 7) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Expected 7 rows x 2 cols, found {num_rows} x {num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table contains correct compensation data (0.30 points)
    # Expected data items (case-insensitive substring matching)
    try:
        expected_items = [
            ("base salary", "$125,000", "125000", "125,000"),
            ("annual bonus", "15%", "15"),
            ("stock options", "5,000 shares", "5000 shares", "4 year"),
            ("health insurance", "80%", "company pays 80"),
            ("401(k)", "6%", "up to 6"),
            ("pto", "20 days", "20"),
        ]

        # Collect all cell text from data rows (skip header)
        data_rows_text = []
        for ri in range(1, len(table.rows)):
            row_text = " ".join(cell.text.strip().lower() for cell in table.rows[ri].cells)
            data_rows_text.append(row_text)

        items_found = 0
        for item_name, *possible_values in expected_items:
            found_in_any_row = False
            for row_text in data_rows_text:
                if item_name in row_text:
                    # Check that at least one expected value substring is present
                    for val in possible_values:
                        if val.lower() in row_text:
                            found_in_any_row = True
                            break
                    if found_in_any_row:
                        break
            if found_in_any_row:
                items_found += 1

        if items_found == 6:
            print(f"PASS: Component 2 — All 6 compensation items found with correct values (0.30 pts)")
            total_score += 0.30
        elif items_found >= 4:
            partial = round(0.30 * (items_found / 6), 2)
            print(f"PARTIAL: Component 2 — {items_found}/6 items found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {items_found}/6 compensation items found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row has gray/dark shading (0.20 points)
    try:
        header_fills = []
        for cell in table.rows[0].cells:
            fill = get_cell_fill(cell)
            header_fills.append(fill)

        # Also check that header text matches expected labels
        header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
        has_header_labels = ("compensation" in header_texts[0] or "component" in header_texts[0]) if len(header_texts) > 0 else False

        all_dark = all(is_dark_gray_fill(f) for f in header_fills)
        if all_dark and has_header_labels:
            print(f"PASS: Component 3 — Header row has dark gray shading (fills={header_fills}) and correct labels (0.20 pts)")
            total_score += 0.20
        elif all_dark:
            print(f"PASS: Component 3 — Header row has dark gray shading (fills={header_fills}) (0.20 pts)")
            total_score += 0.20
        else:
            # Check if any shading at all
            any_fill = any(f is not None and f != 'auto' for f in header_fills)
            if any_fill:
                print(f"PARTIAL: Component 3 — Header has some fill ({header_fills}) but not dark gray (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 3 — Header row has no shading (fills={header_fills})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Alternating row shading on data rows (0.20 points)
    try:
        # Check data rows (rows 1-6) for alternating shading pattern
        # Expected: alternating between no-fill and light-gray fill
        shaded_rows = 0
        unshaded_rows = 0
        pattern_correct = True

        for ri in range(1, len(table.rows)):
            fill = get_cell_fill(table.rows[ri].cells[0])
            if ri % 2 == 0:
                # Even data rows (2, 4, 6) should have light gray shading
                if is_light_gray_fill(fill):
                    shaded_rows += 1
                else:
                    pattern_correct = False
            else:
                # Odd data rows (1, 3, 5) should have no fill or auto
                if fill is None or fill == 'auto':
                    unshaded_rows += 1
                else:
                    pattern_correct = False

        if pattern_correct and shaded_rows >= 2 and unshaded_rows >= 2:
            print(f"PASS: Component 4 — Alternating row shading correct ({shaded_rows} shaded, {unshaded_rows} unshaded) (0.20 pts)")
            total_score += 0.20
        elif shaded_rows >= 1:
            print(f"PARTIAL: Component 4 — Some alternating shading detected but pattern incomplete (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — No alternating row shading detected")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
