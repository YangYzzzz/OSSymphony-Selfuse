"""
Reward Script: Metropolitan Tennis Club Membership Application
Task ID: writer_wf_056
Domain: libreoffice_writer
Scoring:
  C1: Title "Metropolitan Tennis Club" bold 18pt (0.15)
  C2: Applicant info table with 6 fields (0.20)
  C3: Four membership types listed (0.20)
  C4: Sponsor information table with 2 sponsors (0.15)
  C5: Declaration paragraph (0.10)
  C6: Signature and Date lines (0.10)
  C7: Club address in footer (0.10)
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_056'


def persist_app_state(domain: str):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph texts for reuse
    all_texts = [p.text.strip() for p in doc.paragraphs]
    all_texts_lower = [t.lower() for t in all_texts]

    # Component 1: Title "Metropolitan Tennis Club" in bold 18pt (0.15 points)
    try:
        found_title = False
        for para in doc.paragraphs:
            if 'metropolitan tennis club' in para.text.lower():
                # Check for bold and 18pt in at least one run
                for run in para.runs:
                    if run.bold and run.font.size is not None:
                        size_pt = run.font.size.pt
                        if abs(size_pt - 18.0) < 0.5:
                            found_title = True
                            break
                if found_title:
                    break
        if found_title:
            print(f"PASS: Component 1 — Title 'Metropolitan Tennis Club' bold 18pt found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title 'Metropolitan Tennis Club' bold 18pt not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Applicant info table with 6 required fields (0.20 points)
    try:
        required_fields = ['name', 'address', 'phone', 'email', 'date of birth', 'occupation']
        # Also accept common abbreviations
        alt_fields = {
            'date of birth': ['date of birth', 'dob', 'd.o.b', 'birth date', 'birthdate'],
            'phone': ['phone', 'telephone', 'tel', 'phone number', 'contact number'],
            'email': ['email', 'e-mail', 'email address'],
        }

        found_fields = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip().lower()
                    for field in required_fields:
                        if field in alt_fields:
                            if any(alt in cell_text for alt in alt_fields[field]):
                                found_fields.add(field)
                        elif field in cell_text:
                            found_fields.add(field)

        matched = len(found_fields)
        if matched >= 6:
            print(f"PASS: Component 2 — All 6 applicant fields found: {found_fields} (0.20 pts)")
            total_score += 0.20
        elif matched >= 4:
            partial = round(0.20 * matched / 6, 2)
            print(f"PARTIAL: Component 2 — {matched}/6 fields found: {found_fields} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched}/6 applicant fields found: {found_fields}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Four membership types with descriptions (0.20 points)
    try:
        membership_types = ['individual', 'family', 'junior', 'senior']
        found_types = set()
        full_text = ' '.join(all_texts_lower)
        for mtype in membership_types:
            if mtype in full_text:
                found_types.add(mtype)

        matched = len(found_types)
        if matched == 4:
            print(f"PASS: Component 3 — All 4 membership types found: {found_types} (0.20 pts)")
            total_score += 0.20
        elif matched >= 2:
            partial = round(0.20 * matched / 4, 2)
            print(f"PARTIAL: Component 3 — {matched}/4 types found: {found_types} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {matched}/4 membership types found: {found_types}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Sponsor information with 2 sponsors (0.15 points)
    try:
        # Check for sponsor table or sponsor text
        sponsor_count = 0
        # Check tables for sponsor rows
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_lower = cell.text.strip().lower()
                    if 'sponsor 1' in cell_lower or 'sponsor one' in cell_lower or 'first sponsor' in cell_lower:
                        sponsor_count = max(sponsor_count, 1)
                    if 'sponsor 2' in cell_lower or 'sponsor two' in cell_lower or 'second sponsor' in cell_lower:
                        sponsor_count = max(sponsor_count, 2)

        # Also check paragraph text for sponsor references
        if sponsor_count < 2:
            for t in all_texts_lower:
                if 'sponsor 1' in t or 'sponsor one' in t or 'first sponsor' in t:
                    sponsor_count = max(sponsor_count, 1)
                if 'sponsor 2' in t or 'sponsor two' in t or 'second sponsor' in t:
                    sponsor_count = max(sponsor_count, 2)

        # Check if there's a membership number column in sponsor table
        has_membership_num = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if 'membership' in cell.text.strip().lower() and 'number' in cell.text.strip().lower():
                        has_membership_num = True
                        break

        if sponsor_count >= 2 and has_membership_num:
            print(f"PASS: Component 4 — 2 sponsors with membership number fields found (0.15 pts)")
            total_score += 0.15
        elif sponsor_count >= 2:
            print(f"PARTIAL: Component 4 — 2 sponsors found but no membership number field (0.10 pts)")
            total_score += 0.10
        elif sponsor_count >= 1:
            print(f"PARTIAL: Component 4 — Only 1 sponsor found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — No sponsor information found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Declaration paragraph (0.10 points)
    try:
        has_declaration = False
        for t in all_texts_lower:
            if 'declaration' in t or 'hereby' in t or 'agree to abide' in t:
                has_declaration = True
                break
        if has_declaration:
            print(f"PASS: Component 5 — Declaration paragraph found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — No declaration paragraph found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Signature and Date lines (0.10 points)
    try:
        has_signature = False
        has_date = False
        for t in all_texts_lower:
            if 'signature' in t:
                has_signature = True
            if t.startswith('date') and ('_' in t or ':' in t):
                has_date = True

        if has_signature and has_date:
            print(f"PASS: Component 6 — Signature and Date lines found (0.10 pts)")
            total_score += 0.10
        elif has_signature or has_date:
            print(f"PARTIAL: Component 6 — Only {'signature' if has_signature else 'date'} line found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No signature or date lines found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Club address in footer (0.10 points)
    try:
        footer_text = ''
        for section in doc.sections:
            if section.footer and section.footer.paragraphs:
                for fp in section.footer.paragraphs:
                    footer_text += fp.text.strip().lower()

        if 'metropolitan tennis club' in footer_text or ('tennis' in footer_text and 'club' in footer_text):
            print(f"PASS: Component 7 — Club address in footer found (0.10 pts)")
            total_score += 0.10
        elif footer_text and len(footer_text) > 10:
            # Some address-like content in footer
            print(f"PARTIAL: Component 7 — Footer has content but may not have club address: '{footer_text[:80]}' (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 — No club address in footer (footer text: '{footer_text}')")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
