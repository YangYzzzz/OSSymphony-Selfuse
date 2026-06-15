"""
FINAL REWARD SCRIPT - SUCCESS
Task: Whenever I just hammer the Enter key to put some space between my heading and the first paragraph, the layout shifts if I tweak anything later. In LibreOffice Writer, what’s the proper one-step way to insert a horizontal rule directly above Paragraph 1 (so it always sits right between the heading and that very first paragraph, no matter how the text moves)?
Generated: 2025-09-10 20:45:02
Status: success
Model: azure-o3
Total Steps: 6
"""

from docx import Document
from lxml import etree
import os


def has_bottom_border(paragraph):
    """Return True if the given paragraph contains a bottom border (horizontal rule)."""
    # Parse paragraph XML
    xml_str = paragraph._p.xml  # low-level XML string
    root = etree.fromstring(xml_str)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Look for <w:bottom> inside <w:pPr><w:pBdr>
    bottom_elems = root.xpath(".//w:pPr/w:pBdr/w:bottom", namespaces=ns)
    if not bottom_elems:
        return False

    bottom = bottom_elems[0]
    val = bottom.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    sz = bottom.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz")

    # A real rule has val ≠ none and (no size or positive size)
    return bool(val and val.lower() != "none" and (sz is None or int(sz) > 0))


def verify_task(file_path):
    """Verify the LibreOffice Writer task and return a progressive score [0.0-1.0]."""
    print(f"Verifying file: {file_path}")

    max_score = 1.0
    score = 0.0  # progressive score

    # ---------- Prerequisite: file must exist and be loadable (no points) ----------
    if not os.path.exists(file_path):
        print("✗ File not found")
        return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Error loading DOCX: {e}")
        return 0.0

    # ---------- Requirement 1: Heading paragraph present ----------
    heading_para = None
    heading_idx = None
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name.lower().startswith("heading"):
            heading_para = para
            heading_idx = idx
            break

    if heading_para is None:
        print("✗ No heading paragraph detected")
    else:
        print(f"✓ Heading detected: '{heading_para.text.strip()}' (0.3 points)")
        score += 0.3

        # ---------- Requirement 2: Heading has horizontal rule (bottom border) ----------
        if has_bottom_border(heading_para):
            print("✓ Heading has a bottom border (horizontal rule) (0.4 points)")
            score += 0.4
        else:
            print("✗ No bottom border found on heading")

        # ---------- Requirement 3: No blank Enter-spacer paragraphs ----------
        first_content_idx = None
        for idx2 in range(heading_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[idx2].text.strip():  # first non-empty paragraph
                first_content_idx = idx2
                break

        if first_content_idx is None:
            print("✗ No content paragraphs after heading")
        else:
            blank_count = first_content_idx - heading_idx - 1
            if blank_count == 0:
                print("✓ First content paragraph immediately follows heading (0.3 points)")
                score += 0.3
            else:
                print(f"✗ {blank_count} blank paragraph(s) found between heading and first content paragraph")

    final_score = min(score, max_score)
    print(f"Total score: {final_score}/{max_score}")
    return final_score


# ---------------- Execute verification when run as a script ----------------
if __name__ == "__main__":
    file_path = "/home/user/whenever_i_just_hammer_the_enter_key_to_put_some_space_between_my_heading_and_the_first_paragraph_th.docx"
    reward = verify_task(file_path)
    print(f"REWARD: {reward}")
