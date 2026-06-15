"""
Initial Setup: Machine learning survey with 3 bibliography entries (journal articles).
Task ID: writer_bs_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('A Survey of Modern Machine Learning Techniques', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author info ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Dr. Elena Vasquez, Department of Computer Science')
    run.font.size = Pt(11)
    run.italic = True

    # --- Abstract ---
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        'Machine learning has become a cornerstone of modern artificial intelligence, '
        'enabling systems to learn from data and make predictions without explicit programming. '
        'This survey provides an overview of key developments in supervised learning, '
        'unsupervised learning, and reinforcement learning, with particular attention to '
        'neural network architectures that have driven recent advances. We examine foundational '
        'methods and their applications across domains including natural language processing, '
        'computer vision, and autonomous systems.'
    )

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        'The field of machine learning has experienced unprecedented growth over the past '
        'decade. From early statistical methods to modern deep learning approaches, researchers '
        'have developed increasingly sophisticated algorithms capable of tackling complex tasks. '
        'The availability of large datasets and powerful computing resources has accelerated '
        'this progress, leading to breakthroughs in areas such as image recognition, machine '
        'translation, and game-playing agents.'
    )
    doc.add_paragraph(
        'This survey aims to provide a comprehensive overview of the most influential '
        'techniques in contemporary machine learning. We focus on methods that have demonstrated '
        'strong empirical performance and have been widely adopted in both academic research '
        'and industrial applications. Our discussion spans three major paradigms: supervised '
        'learning, unsupervised learning, and reinforcement learning.'
    )

    # --- Section 2: Supervised Learning ---
    doc.add_heading('2. Supervised Learning', level=2)
    doc.add_paragraph(
        'Supervised learning remains the most widely used paradigm in machine learning. '
        'Given labeled training data, algorithms learn a mapping from inputs to outputs that '
        'generalizes to unseen examples. Support vector machines (SVMs) and random forests '
        'have long served as reliable baselines, while deep neural networks have pushed the '
        'boundaries of what is achievable in tasks like object detection and sentiment analysis.'
    )
    doc.add_paragraph(
        'Convolutional neural networks (CNNs) have been particularly transformative for '
        'computer vision tasks. The introduction of architectures such as ResNet and '
        'Inception demonstrated that very deep networks could be trained effectively using '
        'skip connections and batch normalization. Transfer learning from pre-trained models '
        'has further democratized access to powerful vision systems, enabling practitioners '
        'to achieve strong results with limited labeled data.'
    )

    # --- Section 3: Unsupervised Learning ---
    doc.add_heading('3. Unsupervised Learning', level=2)
    doc.add_paragraph(
        'Unsupervised learning algorithms discover structure in unlabeled data. Clustering '
        'methods like k-means and DBSCAN remain fundamental tools for data exploration, while '
        'dimensionality reduction techniques such as PCA and t-SNE facilitate visualization '
        'of high-dimensional datasets. Generative models, including variational autoencoders '
        '(VAEs) and generative adversarial networks (GANs), have opened new possibilities '
        'for data synthesis and augmentation.'
    )

    # --- Section 4: Reinforcement Learning ---
    doc.add_heading('4. Reinforcement Learning', level=2)
    doc.add_paragraph(
        'Reinforcement learning (RL) addresses sequential decision-making problems where '
        'an agent interacts with an environment to maximize cumulative reward. Recent advances '
        'in deep reinforcement learning, combining neural networks with RL algorithms, have '
        'achieved remarkable results. Notable achievements include AlphaGo\'s victory in Go, '
        'robotic manipulation tasks, and autonomous driving systems.'
    )

    # --- Section 5: Conclusion ---
    doc.add_heading('5. Conclusion', level=2)
    doc.add_paragraph(
        'Machine learning continues to evolve rapidly, with new architectures and training '
        'paradigms emerging regularly. The integration of attention mechanisms, transformer '
        'architectures, and self-supervised pre-training has led to powerful foundation models '
        'that excel across multiple modalities. Future research directions include improving '
        'sample efficiency, enhancing interpretability, and developing more robust methods '
        'that perform reliably under distribution shifts.'
    )

    # --- Bibliography (3 journal article entries) ---
    doc.add_heading('Bibliography', level=2)

    bib_entries = [
        (
            '[1] LeCun, Y., Bengio, Y., and Hinton, G. (2015). "Deep Learning." '
            'Nature, Vol. 521, pp. 436-444.'
        ),
        (
            '[2] Krizhevsky, A., Sutskever, I., and Hinton, G. E. (2012). "ImageNet '
            'Classification with Deep Convolutional Neural Networks." Advances in Neural '
            'Information Processing Systems, Vol. 25, pp. 1097-1105.'
        ),
        (
            '[3] Silver, D., Huang, A., Maddison, C. J. et al. (2016). "Mastering the '
            'Game of Go with Deep Neural Networks and Tree Search." Nature, Vol. 529, '
            'pp. 484-489.'
        ),
    ]

    for entry in bib_entries:
        para = doc.add_paragraph(entry)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
