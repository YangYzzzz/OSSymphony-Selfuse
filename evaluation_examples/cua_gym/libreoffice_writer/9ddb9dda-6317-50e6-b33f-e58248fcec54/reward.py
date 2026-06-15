"""
Reward Script: Meeting Minutes Template
Task ID: writer_biz_033
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): 'Meeting Minutes' heading exists with Heading 1 style
  Component 2 (0.3): Table exists with correct dimensions (7 rows x 4 cols)
  Component 3 (0.4): Table header row has correct column names
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_033'

EXPECTED_HEADERS = ['Agenda Item', 'Discussion', 'Action Owner', 'Deadline']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 'Meeting Minutes' in Heading 1 style (0.3 points)
    # This checks that a paragraph with text 'Meeting Minutes' exists and uses Heading 1 style.
    # Initial file is blank, so this will fail on initial_env.
    try:
        heading_matches = [
            p for p in doc.paragraphs
            if p.text.strip() == 'Meeting Minutes' and p.style.name == 'Heading 1'
        ]
        if len(heading_matches) > 0:
            print(f"PASS: Component 1 — 'Meeting Minutes' heading in Heading 1 style (0.3 pts)")
            total_score += 0.3
        else:
            # Check partial: heading text present but wrong style
            has_text = any(p.text.strip() == 'Meeting Minutes' for p in doc.paragraphs)
            has_h1 = any(p.style.name == 'Heading 1' for p in doc.paragraphs)
            print(f"FAIL: Component 1 — heading_text_present={has_text}, heading1_style_present={has_h1}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table with 7 rows and 4 columns (0.3 points)
    # Initial file has no tables, so this fails on initial_env.
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 7 and num_cols == 4:
                print(f"PASS: Component 2 — Table has 7 rows x 4 cols (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Expected 7x4 table, found {num_rows}x{num_cols}")
        else:
            print(f"FAIL: Component 2 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row has correct column names (0.4 points)
    # Checks that row 0 of the first table contains the 4 expected headers.
    # Initial file has no tables, so this fails on initial_env.
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            header_row = table.rows[0]
            actual_headers = [cell.text.strip() for cell in header_row.cells]
            matching = sum(1 for a, e in zip(actual_headers, EXPECTED_HEADERS) if a == e)
            if actual_headers == EXPECTED_HEADERS:
                print(f"PASS: Component 3 — All 4 headers correct: {actual_headers} (0.4 pts)")
                total_score += 0.4
            elif matching >= 2:
                partial = round(0.4 * matching / 4, 2)
                print(f"PARTIAL: Component 3 — {matching}/4 headers correct: {actual_headers} ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Expected {EXPECTED_HEADERS}, found {actual_headers}")
        else:
            print(f"FAIL: Component 3 — No tables found, cannot check headers")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved GUI state before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
