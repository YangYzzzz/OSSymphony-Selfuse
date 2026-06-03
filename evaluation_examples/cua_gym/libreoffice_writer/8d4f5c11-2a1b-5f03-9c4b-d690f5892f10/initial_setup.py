"""
Initial Setup: Create a textbook document with 4 chapters, sub-sections, and a master TOC.
Task ID: writer_mt_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title page ---
    title = doc.add_heading('Fundamentals of Data Science', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Comprehensive Textbook')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Dr. Elena Vasquez & Prof. James Whitmore')
    run.font.size = Pt(13)
    run.italic = True

    edition = doc.add_paragraph()
    edition.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = edition.add_run('Third Edition, 2025')
    run.font.size = Pt(11)

    doc.add_page_break()

    # --- Master Table of Contents ---
    toc_heading = doc.add_heading('Table of Contents', level=1)

    # Chapter structure definition
    chapters = get_chapter_structure()

    # Build master TOC entries
    for ch_idx, chapter in enumerate(chapters, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Chapter {ch_idx}: {chapter["title"]}')
        run.bold = True
        run.font.size = Pt(11)

        for sec in chapter['sections']:
            p = doc.add_paragraph()
            run = p.add_run(f'    {sec["title"]}')
            run.font.size = Pt(10)
            for subsec in sec.get('subsections', []):
                p = doc.add_paragraph()
                run = p.add_run(f'        {subsec}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # --- Chapter Content ---
    for ch_idx, chapter in enumerate(chapters, 1):
        # Chapter heading (Heading 1)
        doc.add_heading(f'Chapter {ch_idx}: {chapter["title"]}', level=1)

        # Chapter intro paragraph
        intro = doc.add_paragraph(chapter['intro'])
        intro.paragraph_format.space_after = Pt(12)

        for sec in chapter['sections']:
            # Section heading (Heading 2)
            doc.add_heading(sec['title'], level=2)
            doc.add_paragraph(sec['content'])

            for subsec_title in sec.get('subsections', []):
                # Sub-section heading (Heading 3)
                doc.add_heading(subsec_title, level=3)
                doc.add_paragraph(sec.get('subsection_content', 'This topic explores key concepts and methodologies that are fundamental to understanding the broader subject area. Researchers have identified several critical factors that influence outcomes in this domain.'))

        # Page break between chapters (except after last)
        if ch_idx < len(chapters):
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def get_chapter_structure():
    return [
        {
            'title': 'Introduction to Data Science',
            'intro': 'Data science has emerged as one of the most impactful disciplines of the 21st century, transforming how organizations make decisions, understand patterns, and predict future outcomes. This chapter provides a foundational overview of the field, its history, and its core principles.',
            'sections': [
                {
                    'title': 'The Evolution of Data Science',
                    'content': 'The roots of data science extend back to the early statistical methods developed in the 18th century. With the advent of computing in the mid-20th century, the ability to process large datasets fundamentally changed how researchers approached analytical problems.',
                    'subsections': ['Historical Milestones', 'The Big Data Revolution'],
                    'subsection_content': 'Key developments in this area have shaped modern analytical practices. From the invention of regression analysis by Francis Galton to the development of machine learning algorithms in the 1990s, each milestone has contributed to the rich tapestry of data science methodology.'
                },
                {
                    'title': 'Core Disciplines and Skill Sets',
                    'content': 'Data science draws from statistics, computer science, and domain expertise. A successful data scientist must be proficient in programming, statistical modeling, and communication of results to stakeholders.',
                    'subsections': ['Statistical Foundations', 'Programming Proficiency'],
                    'subsection_content': 'Mastery of these foundational skills enables practitioners to tackle complex real-world problems effectively. Statistical inference, hypothesis testing, and probability theory form the mathematical backbone of the discipline.'
                },
                {
                    'title': 'The Data Science Lifecycle',
                    'content': 'Every data science project follows a lifecycle from problem definition through data collection, analysis, modeling, and deployment. Understanding this lifecycle is crucial for managing projects effectively.',
                    'subsections': ['Problem Formulation', 'Data Collection Strategies', 'Model Deployment'],
                    'subsection_content': 'Each phase of the lifecycle presents unique challenges and requires specific tools and techniques. Careful planning at each stage ensures reliable and reproducible results.'
                },
                {
                    'title': 'Ethics in Data Science',
                    'content': 'As data science impacts more aspects of daily life, ethical considerations become paramount. Privacy, bias, transparency, and accountability must be addressed in every project.',
                    'subsections': ['Privacy and Consent', 'Algorithmic Bias'],
                    'subsection_content': 'Organizations must establish clear ethical guidelines and review processes to ensure that data science applications do not cause unintended harm to individuals or communities.'
                },
                {
                    'title': 'Industry Applications Overview',
                    'content': 'Data science is applied across healthcare, finance, retail, manufacturing, and government. Each sector presents unique data challenges and opportunities for innovation.',
                    'subsections': ['Healthcare Analytics', 'Financial Modeling'],
                    'subsection_content': 'Sector-specific applications demonstrate the versatility of data science methods. From predicting patient outcomes to detecting fraudulent transactions, the applications are diverse and growing.'
                },
                {
                    'title': 'Tools and Technologies',
                    'content': 'The modern data scientist relies on a toolkit that includes Python, R, SQL, cloud computing platforms, and specialized visualization tools. Choosing the right tool for each task is an essential skill.',
                    'subsections': ['Python Ecosystem', 'Cloud Computing Platforms'],
                    'subsection_content': 'The rapid evolution of tools and platforms means that practitioners must continuously update their skills. Open-source communities have been particularly influential in driving innovation in this space.'
                },
            ]
        },
        {
            'title': 'Statistical Methods and Analysis',
            'intro': 'Statistical methods form the mathematical foundation of data science. This chapter covers essential techniques from descriptive statistics through advanced inferential methods, providing the analytical framework needed for rigorous data analysis.',
            'sections': [
                {
                    'title': 'Descriptive Statistics',
                    'content': 'Descriptive statistics summarize and describe the main features of a dataset. Measures of central tendency, variability, and distribution shape provide the first insights into any dataset.',
                    'subsections': ['Measures of Central Tendency', 'Variability and Spread'],
                    'subsection_content': 'Understanding the center and spread of data distributions is the first step in any analysis. Mean, median, mode, standard deviation, and interquartile range each provide different perspectives on the data.'
                },
                {
                    'title': 'Probability Distributions',
                    'content': 'Probability distributions model the likelihood of different outcomes. The normal, binomial, Poisson, and exponential distributions are among the most commonly used in practice.',
                    'subsections': ['Normal Distribution', 'Discrete Distributions', 'Distribution Fitting'],
                    'subsection_content': 'Selecting the appropriate distribution for modeling a given phenomenon requires understanding the underlying data generation process. Goodness-of-fit tests help validate distributional assumptions.'
                },
                {
                    'title': 'Hypothesis Testing',
                    'content': 'Hypothesis testing provides a formal framework for making decisions based on data. The null and alternative hypotheses, p-values, and significance levels guide the decision-making process.',
                    'subsections': ['Type I and Type II Errors', 'Power Analysis'],
                    'subsection_content': 'Balancing the risk of false positives and false negatives is central to designing effective statistical tests. Sample size calculations ensure adequate power to detect meaningful effects.'
                },
                {
                    'title': 'Regression Analysis',
                    'content': 'Regression models quantify relationships between variables. From simple linear regression to multiple regression with interaction terms, these models are workhorses of predictive analytics.',
                    'subsections': ['Simple Linear Regression', 'Multiple Regression', 'Model Diagnostics'],
                    'subsection_content': 'Interpreting regression coefficients requires attention to confounding variables, multicollinearity, and the assumptions underlying the model. Diagnostic plots reveal potential violations.'
                },
                {
                    'title': 'Bayesian Statistics',
                    'content': 'Bayesian methods offer an alternative framework that incorporates prior knowledge into the analysis. As computational power has increased, Bayesian approaches have become increasingly practical.',
                    'subsections': ['Prior and Posterior Distributions', 'Markov Chain Monte Carlo'],
                    'subsection_content': 'The Bayesian paradigm provides intuitive probabilistic interpretations of results, making it particularly useful in decision-making contexts where prior information is available.'
                },
                {
                    'title': 'Experimental Design',
                    'content': 'Well-designed experiments are essential for establishing causal relationships. Randomization, blocking, and factorial designs help control for confounding variables.',
                    'subsections': ['Randomized Controlled Trials', 'A/B Testing in Practice'],
                    'subsection_content': 'The principles of experimental design apply equally in laboratory settings and digital environments. Online A/B testing has become a standard tool for product development teams.'
                },
                {
                    'title': 'Nonparametric Methods',
                    'content': 'When distributional assumptions cannot be justified, nonparametric methods provide robust alternatives. Rank-based tests and permutation methods are particularly useful for small or skewed samples.',
                    'subsections': ['Rank-Based Tests', 'Bootstrap Methods'],
                    'subsection_content': 'Nonparametric approaches sacrifice some statistical power for increased robustness. The bootstrap method, in particular, has revolutionized inference by enabling estimation of sampling distributions without parametric assumptions.'
                },
            ]
        },
        {
            'title': 'Machine Learning Fundamentals',
            'intro': 'Machine learning automates the process of discovering patterns in data. This chapter introduces the core algorithms, evaluation strategies, and practical considerations that every data scientist must understand to build effective predictive models.',
            'sections': [
                {
                    'title': 'Supervised Learning',
                    'content': 'Supervised learning uses labeled training data to learn a mapping from inputs to outputs. Classification and regression are the two main categories of supervised learning tasks.',
                    'subsections': ['Classification Algorithms', 'Regression Algorithms', 'Feature Engineering'],
                    'subsection_content': 'The choice of algorithm depends on the nature of the data, the size of the training set, and the desired trade-off between interpretability and predictive accuracy. Feature engineering often has a greater impact on performance than algorithm selection.'
                },
                {
                    'title': 'Unsupervised Learning',
                    'content': 'Unsupervised learning discovers hidden structure in unlabeled data. Clustering, dimensionality reduction, and anomaly detection are common unsupervised tasks.',
                    'subsections': ['Clustering Methods', 'Dimensionality Reduction'],
                    'subsection_content': 'Evaluating unsupervised learning results is inherently more challenging than supervised learning since there are no ground truth labels. Internal validation metrics and domain expertise guide the assessment.'
                },
                {
                    'title': 'Model Evaluation and Validation',
                    'content': 'Proper evaluation prevents overfitting and ensures models generalize to new data. Cross-validation, holdout sets, and appropriate metrics are essential components of the evaluation pipeline.',
                    'subsections': ['Cross-Validation Techniques', 'Performance Metrics', 'Overfitting Prevention'],
                    'subsection_content': 'Selecting the right evaluation metric depends on the business context and the relative costs of different types of errors. Accuracy alone is often insufficient, particularly for imbalanced datasets.'
                },
                {
                    'title': 'Ensemble Methods',
                    'content': 'Ensemble methods combine multiple models to achieve better performance than any single model. Random forests, gradient boosting, and stacking are popular ensemble approaches.',
                    'subsections': ['Bagging and Random Forests', 'Gradient Boosting Machines'],
                    'subsection_content': 'The success of ensemble methods is grounded in the diversity of the constituent models. By combining models that make different types of errors, ensembles achieve superior generalization performance.'
                },
                {
                    'title': 'Neural Networks and Deep Learning',
                    'content': 'Neural networks learn hierarchical representations of data through layers of interconnected nodes. Deep learning has achieved remarkable results in image recognition, natural language processing, and beyond.',
                    'subsections': ['Feedforward Networks', 'Convolutional Neural Networks', 'Recurrent Networks'],
                    'subsection_content': 'Training deep networks requires careful attention to architecture design, learning rate scheduling, regularization, and the choice of activation functions. Transfer learning has made deep learning accessible for tasks with limited training data.'
                },
                {
                    'title': 'Practical Machine Learning Pipelines',
                    'content': 'Moving from prototype to production requires robust pipelines that handle data ingestion, preprocessing, training, evaluation, and serving. MLOps practices ensure reliability and reproducibility.',
                    'subsections': ['Data Pipeline Design', 'Model Serving and Monitoring'],
                    'subsection_content': 'Production machine learning systems must handle data drift, model degradation, and changing requirements. Continuous integration and deployment practices adapted for ML help maintain system quality.'
                },
            ]
        },
        {
            'title': 'Data Visualization and Communication',
            'intro': 'The ability to communicate findings effectively is what transforms analysis into actionable insight. This chapter covers visualization principles, tools, and strategies for presenting data to both technical and non-technical audiences.',
            'sections': [
                {
                    'title': 'Principles of Effective Visualization',
                    'content': 'Good visualization is guided by principles of clarity, accuracy, and efficiency. Edward Tufte\'s data-ink ratio and the grammar of graphics provide theoretical foundations for creating effective charts.',
                    'subsections': ['Tufte\'s Principles', 'The Grammar of Graphics'],
                    'subsection_content': 'Understanding the perceptual and cognitive factors that influence how people interpret visual information is essential for creating charts that communicate accurately and efficiently.'
                },
                {
                    'title': 'Chart Types and Selection',
                    'content': 'Choosing the right chart type depends on the data structure and the story you want to tell. Bar charts, line charts, scatter plots, heatmaps, and box plots each serve different purposes.',
                    'subsections': ['Comparison Charts', 'Distribution Charts', 'Relationship Charts'],
                    'subsection_content': 'Each chart type has strengths and limitations. Understanding when to use a histogram versus a density plot, or a grouped bar chart versus a stacked bar chart, is a skill that improves with practice and feedback.'
                },
                {
                    'title': 'Interactive Visualizations',
                    'content': 'Interactive visualizations allow users to explore data dynamically. Dashboards, drill-down charts, and linked views provide richer experiences than static images.',
                    'subsections': ['Dashboard Design', 'User Interaction Patterns'],
                    'subsection_content': 'Effective interactive visualizations balance complexity with usability. Too many interactive elements can overwhelm users, while too few may limit the exploratory value of the visualization.'
                },
                {
                    'title': 'Visualization Tools and Libraries',
                    'content': 'Modern visualization tools range from code-based libraries like matplotlib, seaborn, and plotly to GUI-based tools like Tableau and Power BI. Each has its strengths for different use cases.',
                    'subsections': ['Python Visualization Stack', 'Business Intelligence Tools'],
                    'subsection_content': 'The choice between code-based and GUI-based tools depends on the audience, the need for reproducibility, and the complexity of the visualization. Many practitioners use both approaches in their workflow.'
                },
                {
                    'title': 'Storytelling with Data',
                    'content': 'Data storytelling combines analytical rigor with narrative techniques to drive action. Effective stories have a clear beginning, middle, and end, supported by evidence from the data.',
                    'subsections': ['Narrative Structure', 'Audience Adaptation'],
                    'subsection_content': 'Tailoring the message to the audience is perhaps the most important aspect of data communication. Technical details that engage a data science team may confuse a board of directors, and vice versa.'
                },
            ]
        },
    ]


create_initial()
