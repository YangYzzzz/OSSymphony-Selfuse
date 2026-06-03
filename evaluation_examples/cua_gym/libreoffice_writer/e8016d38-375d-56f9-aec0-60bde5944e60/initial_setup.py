"""
Initial Setup: Create a Terms document with 3 pages of content, no macros.
Task ID: writer_tm_079
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_079'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('Terms of Service Agreement', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Software Solutions, Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph('')  # spacer

    # --- Section 1: Acceptance of Terms ---
    doc.add_heading('1. Acceptance of Terms', level=1)
    doc.add_paragraph(
        'By accessing or using the services provided by Meridian Software Solutions, Inc. '
        '("Company," "we," "us," or "our"), you agree to be bound by these Terms of Service '
        '("Terms"). If you do not agree to all of these Terms, you may not access or use our '
        'services. These Terms apply to all visitors, users, and others who access or use the '
        'services, including but not limited to browsing our website, purchasing subscriptions, '
        'or utilizing our cloud-based software tools.'
    )
    doc.add_paragraph(
        'We reserve the right to update or modify these Terms at any time without prior notice. '
        'Your continued use of our services after any such changes constitutes your acceptance '
        'of the new Terms. It is your responsibility to review these Terms periodically for '
        'updates. The most current version of these Terms will always be available on our website '
        'at www.meridiansoftware.com/terms.'
    )

    # --- Section 2: Description of Services ---
    doc.add_heading('2. Description of Services', level=1)
    doc.add_paragraph(
        'Meridian Software Solutions provides a suite of cloud-based project management and '
        'collaboration tools designed for enterprise teams. Our platform includes task tracking, '
        'document sharing, real-time communication channels, automated workflow builders, and '
        'comprehensive analytics dashboards. The services are offered on a subscription basis '
        'with tiered pricing plans: Starter ($29/month per user), Professional ($59/month per '
        'user), and Enterprise (custom pricing).'
    )
    doc.add_paragraph(
        'Each subscription tier provides different levels of access to features and storage '
        'capacity. The Starter plan includes up to 10 GB of cloud storage per user, while the '
        'Professional plan offers 50 GB and the Enterprise plan provides unlimited storage. '
        'Additional storage can be purchased in increments of 100 GB at $4.99 per month. All '
        'plans include basic customer support via email with a guaranteed response time of 48 '
        'hours for Starter, 24 hours for Professional, and 4 hours for Enterprise.'
    )

    # --- Section 3: User Accounts and Registration ---
    doc.add_heading('3. User Accounts and Registration', level=1)
    doc.add_paragraph(
        'To access most features of our services, you must create an account. When creating an '
        'account, you agree to provide accurate, current, and complete information. You are '
        'responsible for maintaining the confidentiality of your account credentials and for all '
        'activities that occur under your account. You must immediately notify us of any '
        'unauthorized use of your account or any other security breach.'
    )
    doc.add_paragraph(
        'Account holders must be at least 18 years of age or the age of majority in their '
        'jurisdiction, whichever is greater. Corporate accounts must be registered by an '
        'authorized representative of the organization. Each user within a corporate account '
        'must have their own individual login credentials. Sharing of login credentials between '
        'multiple individuals is strictly prohibited and may result in immediate account '
        'termination without refund.'
    )

    # --- Section 4: Payment and Billing ---
    doc.add_heading('4. Payment and Billing', level=1)
    doc.add_paragraph(
        'Subscription fees are billed in advance on a monthly or annual basis, depending on the '
        'billing cycle selected at the time of purchase. Annual subscriptions receive a 20% '
        'discount compared to monthly pricing. All payments are processed through our secure '
        'payment gateway and must be made using a valid credit card, debit card, or approved '
        'electronic payment method. Invoicing is available for Enterprise plan customers with '
        'payment terms of Net 30.'
    )
    doc.add_paragraph(
        'If payment fails, we will attempt to process the charge up to three additional times '
        'over a 10-day period. If all attempts fail, your account will be downgraded to a '
        'read-only state, and you will have 30 days to update your payment information before '
        'your data is archived. Archived data is retained for 90 days, after which it is '
        'permanently deleted. Refunds are available within the first 14 days of a new '
        'subscription or plan upgrade, subject to our Refund Policy.'
    )

    # --- Section 5: Intellectual Property ---
    doc.add_heading('5. Intellectual Property Rights', level=1)
    doc.add_paragraph(
        'All content, features, and functionality of our services, including but not limited to '
        'software code, text, graphics, logos, icons, images, audio clips, data compilations, '
        'and the compilation thereof, are the exclusive property of Meridian Software Solutions, '
        'Inc. and are protected by United States and international copyright, trademark, patent, '
        'trade secret, and other intellectual property laws. Our trademarks and trade dress may '
        'not be used in connection with any product or service that is not ours.'
    )
    doc.add_paragraph(
        'You retain ownership of any content you upload, create, or share through our services. '
        'However, by using our services, you grant us a non-exclusive, worldwide, royalty-free '
        'license to use, reproduce, modify, and display your content solely for the purpose of '
        'providing and improving our services. This license terminates when you delete your '
        'content from our platform or close your account, except where your content has been '
        'shared with other users who have not deleted it.'
    )

    # --- Section 6: Privacy and Data Protection ---
    doc.add_heading('6. Privacy and Data Protection', level=1)
    doc.add_paragraph(
        'Your privacy is critically important to us. Our collection and use of personal '
        'information is governed by our Privacy Policy, which is incorporated into these Terms '
        'by reference. We comply with the General Data Protection Regulation (GDPR), the '
        'California Consumer Privacy Act (CCPA), and other applicable data protection laws. '
        'We implement industry-standard security measures including AES-256 encryption for '
        'data at rest and TLS 1.3 for data in transit.'
    )
    doc.add_paragraph(
        'We process personal data on servers located in the United States, the European Union, '
        'and Singapore. Enterprise customers may select their preferred data residency region. '
        'We conduct annual third-party security audits and maintain SOC 2 Type II certification. '
        'In the event of a data breach affecting your personal information, we will notify you '
        'within 72 hours in compliance with applicable regulations. You may exercise your rights '
        'to access, rectify, or delete your personal data by contacting privacy@meridiansoftware.com.'
    )

    # --- Section 7: Limitation of Liability ---
    doc.add_heading('7. Limitation of Liability', level=1)
    doc.add_paragraph(
        'To the maximum extent permitted by applicable law, Meridian Software Solutions, Inc. '
        'shall not be liable for any indirect, incidental, special, consequential, or punitive '
        'damages, including but not limited to loss of profits, data, use, goodwill, or other '
        'intangible losses, resulting from (a) your access to or use of or inability to access '
        'or use the services; (b) any conduct or content of any third party on the services; '
        '(c) any content obtained from the services; or (d) unauthorized access, use, or '
        'alteration of your transmissions or content.'
    )
    doc.add_paragraph(
        'In no event shall our total aggregate liability exceed the amount you have paid us in '
        'the twelve (12) months preceding the claim, or one hundred dollars ($100), whichever '
        'is greater. This limitation of liability applies whether the alleged liability is based '
        'on contract, tort, negligence, strict liability, or any other basis, even if we have '
        'been advised of the possibility of such damage. Some jurisdictions do not allow the '
        'exclusion or limitation of certain damages, so some of these limitations may not apply '
        'to you.'
    )

    # --- Section 8: Governing Law ---
    doc.add_heading('8. Governing Law and Dispute Resolution', level=1)
    doc.add_paragraph(
        'These Terms shall be governed by and construed in accordance with the laws of the '
        'State of Delaware, United States, without regard to its conflict of law principles. '
        'Any dispute arising out of or relating to these Terms or the services shall first be '
        'submitted to mediation administered by the American Arbitration Association. If '
        'mediation is unsuccessful, the dispute shall be resolved by binding arbitration in '
        'Wilmington, Delaware, conducted in English under the Commercial Arbitration Rules.'
    )
    doc.add_paragraph(
        'Notwithstanding the foregoing, either party may seek injunctive or other equitable '
        'relief in any court of competent jurisdiction to prevent the actual or threatened '
        'infringement or misappropriation of intellectual property rights. Class action lawsuits, '
        'class-wide arbitrations, private attorney-general actions, and any other proceeding '
        'where someone acts in a representative capacity are not permitted. The arbitrator may '
        'award relief only in favor of the individual party seeking relief and only to the '
        'extent necessary to provide relief warranted by that individual party\'s claim.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
