"""
Reward Script: Content Audit Document — Bulleted List to Formatted Table
Task ID: writer_mktg_060
Domain: libreoffice_writer
Scoring:
  Component 1: Table structure (26 rows x 6 cols, correct header names) — 0.25 pts
  Component 2: Header row formatting (bold, white text, #424242 background) — 0.25 pts
  Component 3: Status column color-coding (Current=green, Needs Update=yellow, Retire=red) — 0.30 pts
  Component 4: Validation note (italic, 9pt, correct text) — 0.20 pts
  Total: 1.0
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_060'
FILE_PATH = f'{WORKDIR}/content_audit.docx'

# Expected column headers
EXPECTED_HEADERS = ['Asset Title', 'Type', 'Owner', 'Last Updated', 'Status', 'Priority']

# Expected status color mapping (hex, uppercase)
STATUS_COLORS = {
    'Current': 'E8F5E9',
    'Needs Update': 'FFF9C4',
    'Retire': 'FFEBEE',
}

# Expected header background color
HEADER_BG_COLOR = '424242'

# Validation note keywords
VALIDATION_NOTE_KEYWORDS = ['Status:', 'Current', 'Needs Update', 'Retire', 'Priority:', 'High', 'Medium', 'Low']


def get_cell_fill(cell):
    """Return hex fill color (uppercase) for a table cell, or None."""
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill.upper() not in ('', 'AUTO', 'FFFFFF00'):
                return fill.upper()
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Table structure — 26 rows x 6 cols with correct header names (0.25 pts)
    # The initial file has NO table; golden file has a 26-row x 6-col table.
    # -------------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 — No table found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)

            # Check dimensions: 26 rows (1 header + 25 data) x 6 cols
            dims_ok = (num_rows == 26 and num_cols == 6)
            if not dims_ok:
                print(f"FAIL: Component 1 — Table dimensions {num_rows}x{num_cols}, expected 26x6")
            else:
                # Check header names
                header_cells = [table.rows[0].cells[j].text.strip() for j in range(6)]
                if header_cells == EXPECTED_HEADERS:
                    print(f"PASS: Component 1 — Table 26x6 with correct headers {header_cells} (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 1 — Header names mismatch: found {header_cells}, expected {EXPECTED_HEADERS}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Header row formatting — bold text, white color (#FFFFFF),
    # dark gray background (#424242) (0.25 pts)
    # -------------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No table found")
        else:
            table = doc.tables[0]
            header_row = table.rows[0]
            all_bold = True
            all_white_text = True
            all_dark_bg = True
            issues = []

            for j, cell in enumerate(header_row.cells):
                # Check background
                fill = get_cell_fill(cell)
                if fill != HEADER_BG_COLOR:
                    all_dark_bg = False
                    issues.append(f"Col {j} bg={fill!r} expected {HEADER_BG_COLOR!r}")

                # Check run formatting
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.bold:
                            all_bold = False
                            issues.append(f"Col {j} run not bold: {run.text!r}")
                        try:
                            rgb = run.font.color.rgb
                            if str(rgb).upper() != 'FFFFFF':
                                all_white_text = False
                                issues.append(f"Col {j} text color={rgb}, expected FFFFFF")
                        except Exception:
                            # color.rgb raises if no color set
                            all_white_text = False
                            issues.append(f"Col {j} no explicit text color")

            if all_bold and all_white_text and all_dark_bg:
                print(f"PASS: Component 2 — Header row bold, white text, #424242 background (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — Header formatting issues: {issues[:3]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Status column color-coding (0.30 pts)
    # Current → #E8F5E9 (green), Needs Update → #FFF9C4 (yellow), Retire → #FFEBEE (red)
    # Every data row's Status cell must have the correct background.
    # -------------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No table found")
        else:
            table = doc.tables[0]
            if len(table.rows) < 2:
                print("FAIL: Component 3 — Table has no data rows")
            else:
                correct = 0
                total_data = 0
                fail_examples = []

                for i, row in enumerate(table.rows[1:], 1):
                    status_cell = row.cells[4]
                    status_text = status_cell.text.strip()

                    if status_text not in STATUS_COLORS:
                        # Skip rows with unexpected status values
                        continue

                    total_data += 1
                    expected_fill = STATUS_COLORS[status_text]
                    actual_fill = get_cell_fill(status_cell)

                    if actual_fill == expected_fill:
                        correct += 1
                    else:
                        fail_examples.append(
                            f"Row {i}: status={status_text!r} fill={actual_fill!r} expected={expected_fill!r}"
                        )

                if total_data == 0:
                    print("FAIL: Component 3 — No data rows with known Status values found")
                else:
                    ratio = correct / total_data
                    if ratio == 1.0:
                        print(f"PASS: Component 3 — All {total_data} Status cells correctly color-coded (0.30 pts)")
                        total_score += 0.30
                    elif ratio >= 0.8:
                        # Partial credit: mostly correct
                        partial = round(0.30 * ratio, 2)
                        print(
                            f"PARTIAL: Component 3 — {correct}/{total_data} Status cells correct "
                            f"({partial} pts). Examples of failures: {fail_examples[:3]}"
                        )
                        total_score += partial
                    else:
                        print(
                            f"FAIL: Component 3 — Only {correct}/{total_data} Status cells correctly "
                            f"color-coded. Examples: {fail_examples[:3]}"
                        )
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Validation note paragraph (0.20 pts)
    # Must exist outside the table: italic, 9pt, text containing
    # "Status: Current | Needs Update | Retire. Priority: High | Medium | Low."
    # -------------------------------------------------------------------------
    try:
        note_found = False
        for para in doc.paragraphs:
            text = para.text.strip()
            # Check if all required keywords are present
            if all(kw in text for kw in VALIDATION_NOTE_KEYWORDS):
                # Check formatting: italic and ~9pt
                is_italic = False
                is_9pt = False
                for run in para.runs:
                    if run.text.strip():
                        if run.italic:
                            is_italic = True
                        if run.font.size is not None:
                            size_pt = run.font.size.pt
                            if abs(size_pt - 9.0) < 0.5:
                                is_9pt = True

                if is_italic and is_9pt:
                    print(f"PASS: Component 4 — Validation note found, italic, 9pt: {text[:60]!r} (0.20 pts)")
                    total_score += 0.20
                    note_found = True
                elif is_italic:
                    print(f"PARTIAL: Component 4 — Validation note found, italic but not 9pt (0.10 pts): {text[:60]!r}")
                    total_score += 0.10
                    note_found = True
                else:
                    print(f"FAIL: Component 4 — Validation note text found but not italic/9pt: italic={is_italic}, 9pt={is_9pt}")
                    note_found = True
                break

        if not note_found:
            print("FAIL: Component 4 — No validation note paragraph found with required keywords")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
