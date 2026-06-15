"""
Initial Setup: Marketing Content Audit document with bulleted list
Task ID: writer_mktg_060
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'content_audit'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title
    title_para = doc.add_heading('Marketing Content Audit — Q1 2026', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Introduction paragraph
    intro = doc.add_paragraph(
        'This document provides a comprehensive audit of all current marketing content assets. '
        'Each asset has been reviewed for relevance, accuracy, and alignment with our Q1 2026 objectives. '
        'Use this audit to prioritize content updates, identify gaps, and retire outdated materials.'
    )

    # 25 content asset bullet items
    # Each bullet: "Title | Type | Owner | Last Updated | Status"
    assets = [
        ("2025 Annual Product Brochure", "one-pager", "Emily Torres", "2025-01-10", "Needs Update", "High"),
        ("Cloud Services Whitepaper", "whitepaper", "James Liu", "2024-11-20", "Current", "High"),
        ("Customer Success Story: RetailCo", "case study", "Maria Gonzalez", "2024-09-05", "Current", "Medium"),
        ("Introduction to Our Platform", "video", "Kevin Park", "2024-07-14", "Needs Update", "Medium"),
        ("Email Campaign: Holiday 2024", "blog", "Sarah Chen", "2024-12-01", "Retire", "Low"),
        ("SEO Blog: Top 10 B2B Trends", "blog", "Ashley White", "2025-02-15", "Current", "High"),
        ("Security & Compliance One-Pager", "one-pager", "David Kim", "2024-08-22", "Needs Update", "High"),
        ("Integration Guide: Salesforce", "whitepaper", "Rachel Brown", "2024-10-11", "Current", "Medium"),
        ("Product Demo Video: Enterprise", "video", "Marcus Johnson", "2025-01-25", "Current", "High"),
        ("Customer Webinar: Q4 Recap", "video", "Lisa Nguyen", "2024-12-18", "Retire", "Low"),
        ("Case Study: FinanceCorp ROI", "case study", "Tom Stevens", "2025-02-01", "Current", "High"),
        ("Blog: How AI Is Changing CRM", "blog", "Jessica Park", "2024-06-30", "Retire", "Low"),
        ("Product Comparison One-Pager", "one-pager", "Brian Adams", "2025-01-05", "Current", "Medium"),
        ("Enterprise Pricing Whitepaper", "whitepaper", "Natalie Chen", "2024-11-02", "Needs Update", "High"),
        ("Case Study: HealthNet Savings", "case study", "Omar Hassan", "2024-10-28", "Current", "Medium"),
        ("Blog: 5 Ways to Boost Retention", "blog", "Priya Sharma", "2025-02-20", "Current", "Medium"),
        ("Onboarding Video Series Ep.1", "video", "Daniel Wright", "2024-08-10", "Needs Update", "High"),
        ("Partner Program Overview", "one-pager", "Chloe Martin", "2024-07-03", "Retire", "Low"),
        ("API Documentation Guide", "whitepaper", "Ethan Clark", "2025-01-18", "Current", "High"),
        ("Blog: Remote Work Best Practices", "blog", "Mia Robinson", "2024-05-15", "Retire", "Low"),
        ("Case Study: LogisticsHub Speed", "case study", "Aiden Lee", "2025-02-10", "Current", "Medium"),
        ("Event Recap: TechSummit 2024", "blog", "Sofia Walker", "2024-11-15", "Needs Update", "Medium"),
        ("Mobile App Feature Overview", "one-pager", "Noah Thompson", "2025-01-28", "Current", "High"),
        ("Q1 2025 State of SaaS Report", "whitepaper", "Isabella Davis", "2025-03-01", "Current", "High"),
        ("Video Testimonial: StartupXYZ", "video", "Liam Martinez", "2024-09-20", "Needs Update", "Medium"),
    ]

    for asset in assets:
        title, atype, owner, last_updated, status, priority = asset
        bullet_text = (
            f"{title} | Type: {atype} | Owner: {owner} | "
            f"Last Updated: {last_updated} | Status: {status} | Priority: {priority}"
        )
        doc.add_paragraph(bullet_text, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
