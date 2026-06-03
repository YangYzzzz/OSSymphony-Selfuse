"""
Reward Script: Certificate of Appreciation formatting task
Task ID: writer_creative_043
Domain: libreoffice_writer
Scoring:
  - Component 1: Page orientation is landscape (0.20 pts)
  - Component 2: Double-line page border on all 4 sides (0.20 pts)
  - Component 3: Title 'Certificate of Appreciation' is 32pt, bold, center, space_before=72pt (0.20 pts)
  - Component 4: 'Karen Mitchell' is 28pt, bold, italic, center (0.15 pts)
  - Component 5: 'Presented to' is 16pt, center, space_before=36pt (0.10 pts)
  - Component 6: Description line is 14pt, center, space_before=24pt (0.10 pts)
  - Component 7: All paragraphs use Liberation Serif font family (0.05 pts)
  Total: 1.00
"""

import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'volunteer_certificate'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: Page orientation is landscape (0.20 points)
    # This should FAIL on initial (portrait) and PASS on golden (landscape)
    try:
        is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
        if is_landscape:
            print(f"PASS: Component 1 — Page orientation is LANDSCAPE (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected LANDSCAPE orientation, found: {section.orientation}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Double-line page border on all 4 sides (0.20 points)
    # This should FAIL on initial (no border) and PASS on golden (double border)
    try:
        sectPr = section._sectPr
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        pgBorders = sectPr.find(f'{{{ns}}}pgBorders')

        if pgBorders is not None:
            sides = ['top', 'left', 'bottom', 'right']
            missing_sides = []

            for side in sides:
                elem = pgBorders.find(f'{{{ns}}}{side}')
                if elem is None:
                    missing_sides.append(side)
                else:
                    val = elem.get(f'{{{ns}}}val')
                    if val != 'double':
                        missing_sides.append(f"{side}(val={val})")

            if not missing_sides:
                print(f"PASS: Component 2 — Double-line border on all 4 sides (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — Border not double on sides: {missing_sides}")
        else:
            print(f"FAIL: Component 2 — No page border element found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Build paragraph lookup by text content
    para_map = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            para_map[text] = para

    # Component 3: Title 'Certificate of Appreciation' is 32pt, bold, center, space_before=72pt (0.20 pts)
    # This should FAIL on initial (12pt, not bold, left) and PASS on golden
    try:
        title_text = 'Certificate of Appreciation'
        title_para = None
        for key, para in para_map.items():
            if title_text in key:
                title_para = para
                break

        if title_para is None:
            print(f"FAIL: Component 3 — 'Certificate of Appreciation' paragraph not found")
        else:
            pf = title_para.paragraph_format
            runs = title_para.runs

            checks = []
            # Check alignment
            if pf.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                checks.append('alignment=CENTER')
            else:
                print(f"  FAIL detail: alignment={pf.alignment}, expected CENTER")

            # Check space_before (72pt)
            sb = pf.space_before
            if sb is not None and abs(sb.pt - 72.0) < 1.0:
                checks.append('space_before=72pt')
            else:
                print(f"  FAIL detail: space_before={sb.pt if sb else None}pt, expected 72pt")

            # Check first run size and bold
            if runs:
                run = runs[0]
                if run.font.size is not None and abs(run.font.size.pt - 32.0) < 0.5:
                    checks.append('size=32pt')
                else:
                    print(f"  FAIL detail: size={run.font.size.pt if run.font.size else None}pt, expected 32pt")

                if run.font.bold:
                    checks.append('bold=True')
                else:
                    print(f"  FAIL detail: bold={run.font.bold}, expected True")

            if len(checks) == 4:
                print(f"PASS: Component 3 — Title is 32pt, bold, CENTER, space_before=72pt (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Only {len(checks)}/4 title checks passed: {checks}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 'Karen Mitchell' is 28pt, bold, italic, center (0.15 points)
    # This should FAIL on initial (12pt, not bold/italic, left) and PASS on golden
    try:
        km_text = 'Karen Mitchell'
        km_para = None
        for key, para in para_map.items():
            if km_text in key:
                km_para = para
                break

        if km_para is None:
            print(f"FAIL: Component 4 — 'Karen Mitchell' paragraph not found")
        else:
            pf = km_para.paragraph_format
            runs = km_para.runs

            checks = []
            if pf.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                checks.append('alignment=CENTER')
            else:
                print(f"  FAIL detail: alignment={pf.alignment}, expected CENTER")

            if runs:
                run = runs[0]
                if run.font.size is not None and abs(run.font.size.pt - 28.0) < 0.5:
                    checks.append('size=28pt')
                else:
                    print(f"  FAIL detail: size={run.font.size.pt if run.font.size else None}pt, expected 28pt")

                if run.font.bold:
                    checks.append('bold=True')
                else:
                    print(f"  FAIL detail: bold={run.font.bold}, expected True")

                if run.font.italic:
                    checks.append('italic=True')
                else:
                    print(f"  FAIL detail: italic={run.font.italic}, expected True")

            if len(checks) == 4:
                print(f"PASS: Component 4 — 'Karen Mitchell' is 28pt, bold, italic, CENTER (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — Only {len(checks)}/4 Karen Mitchell checks passed: {checks}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 'Presented to' is 16pt, center, space_before=36pt (0.10 points)
    # This should FAIL on initial (12pt, left, no space) and PASS on golden
    try:
        pt_text = 'Presented to'
        pt_para = None
        for key, para in para_map.items():
            if pt_text in key:
                pt_para = para
                break

        if pt_para is None:
            print(f"FAIL: Component 5 — 'Presented to' paragraph not found")
        else:
            pf = pt_para.paragraph_format
            runs = pt_para.runs

            checks = []
            if pf.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                checks.append('alignment=CENTER')
            else:
                print(f"  FAIL detail: alignment={pf.alignment}, expected CENTER")

            sb = pf.space_before
            if sb is not None and abs(sb.pt - 36.0) < 1.0:
                checks.append('space_before=36pt')
            else:
                print(f"  FAIL detail: space_before={sb.pt if sb else None}pt, expected 36pt")

            if runs:
                run = runs[0]
                if run.font.size is not None and abs(run.font.size.pt - 16.0) < 0.5:
                    checks.append('size=16pt')
                else:
                    print(f"  FAIL detail: size={run.font.size.pt if run.font.size else None}pt, expected 16pt")

            if len(checks) == 3:
                print(f"PASS: Component 5 — 'Presented to' is 16pt, CENTER, space_before=36pt (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 5 — Only {len(checks)}/3 'Presented to' checks passed: {checks}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Description line is 14pt, center, space_before=24pt (0.10 points)
    # This should FAIL on initial (12pt, left, no space) and PASS on golden
    try:
        desc_keywords = ['In recognition of']
        desc_para = None
        for key, para in para_map.items():
            for kw in desc_keywords:
                if kw in key:
                    desc_para = para
                    break
            if desc_para:
                break

        if desc_para is None:
            print(f"FAIL: Component 6 — Description paragraph not found")
        else:
            pf = desc_para.paragraph_format
            runs = desc_para.runs

            checks = []
            if pf.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                checks.append('alignment=CENTER')
            else:
                print(f"  FAIL detail: alignment={pf.alignment}, expected CENTER")

            sb = pf.space_before
            if sb is not None and abs(sb.pt - 24.0) < 1.0:
                checks.append('space_before=24pt')
            else:
                print(f"  FAIL detail: space_before={sb.pt if sb else None}pt, expected 24pt")

            if runs:
                run = runs[0]
                if run.font.size is not None and abs(run.font.size.pt - 14.0) < 0.5:
                    checks.append('size=14pt')
                else:
                    print(f"  FAIL detail: size={run.font.size.pt if run.font.size else None}pt, expected 14pt")

            if len(checks) == 3:
                print(f"PASS: Component 6 — Description is 14pt, CENTER, space_before=24pt (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 6 — Only {len(checks)}/3 description checks passed: {checks}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: All non-empty paragraphs use Liberation Serif font family (0.05 points)
    # This should FAIL on initial (Times New Roman) and PASS on golden (Liberation Serif)
    try:
        expected_font = 'Liberation Serif'
        wrong_fonts = []
        for para in doc.paragraphs:
            if para.text.strip():
                for run in para.runs:
                    if run.text.strip() and run.font.name and run.font.name != expected_font:
                        wrong_fonts.append(f"'{run.text[:20]}': {run.font.name}")

        if not wrong_fonts:
            print(f"PASS: Component 7 — All text uses Liberation Serif font (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 — Wrong fonts found: {wrong_fonts[:3]}")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
