"""
Initial Setup: Certificate of Appreciation document (pre-task state)
Task ID: writer_creative_043
Domain: libreoffice_writer

Creates: /home/user/volunteer_certificate.docx
  - Contains all certificate text in plain, unformatted style
  - 12pt Times New Roman, left-aligned, portrait orientation
  - No page border, no large fonts, no special formatting
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'volunteer_certificate'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set portrait orientation (default)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    # Standard letter size: 8.5" x 11"
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Define all certificate lines — plain 12pt Times New Roman, left-aligned
    lines = [
        'Certificate of Appreciation',
        'Presented to',
        'Karen Mitchell',
        'In recognition of 5 years of dedicated service to the Riverside Food Bank',
        'March 4, 2026',
        'Director\'s Signature: _____________________',
    ]

    for line_text in lines:
        para = doc.add_paragraph()
        run = para.add_run(line_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = False
        run.italic = False
        # Left-aligned (default), no special spacing
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
