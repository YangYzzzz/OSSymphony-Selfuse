"""
Initial Setup: Export the open document as a PDF, preserving the base filename and output directory.
Task ID: osworld_writer_pdf_export_keepname_009
Domain: libreoffice_writer

Creates:
  /home/user/Reports/quarterly_review.odt   (source document, NO pdf present)

Then opens the .odt file in LibreOffice Writer for the GUI agent.
"""

import os
import shlex
import subprocess
import time

# -------------------------------------------------------------------
# Paths (VM-side)
# -------------------------------------------------------------------
WORKDIR = '/home/user'
REPORTS_DIR = '/home/user/Reports'
TASK_ID = 'quarterly_review'
ODT_FILE = f'{REPORTS_DIR}/{TASK_ID}.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch a GUI app on the VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # 1. Ensure the Reports directory exists
    os.makedirs(REPORTS_DIR, exist_ok=True)

    # 2. Remove any stale PDF that might already exist (keep initial clean)
    pdf_path = f'{REPORTS_DIR}/{TASK_ID}.pdf'
    if os.path.exists(pdf_path):
        os.remove(pdf_path)

    # 3. Build a realistic quarterly review ODT document using python-docx
    #    then save as .odt via the python-docx-compatible approach.
    #    Because python-docx natively writes .docx (OOXML), we write a .docx
    #    then convert it to .odt with LibreOffice headlessly.

    import tempfile
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # --- Cover / Title Section ---
    title_para = doc.add_heading('Acme Corporation', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub_para = doc.add_paragraph('Q3 2024 — Quarterly Business Review')
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in sub_para.runs:
        run.font.size = Pt(14)
        run.font.italic = True

    date_para = doc.add_paragraph('Prepared by: Finance & Strategy Team  |  September 30, 2024')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in date_para.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()  # spacer

    # --- Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    exec_summary = (
        'Acme Corporation delivered strong results in Q3 2024, achieving $128.4 million in total revenue, '
        'representing a 14.2% year-over-year increase. Operating income reached $31.6 million (24.6% margin), '
        'driven by robust demand in the Enterprise segment and disciplined cost management across all business units. '
        'Free cash flow improved to $22.1 million, enabling continued investment in R&D and strategic acquisitions.'
    )
    doc.add_paragraph(exec_summary)

    highlights_intro = doc.add_paragraph()
    highlights_intro.add_run('Key Highlights:').bold = True

    highlights = [
        'Revenue: $128.4M (+14.2% YoY)',
        'Gross Margin: 61.3% (vs. 59.8% prior year)',
        'Operating Income: $31.6M (24.6% margin)',
        'Net Income: $23.9M (+18.7% YoY)',
        'Customer Retention Rate: 94.2%',
        'New Enterprise Logos: 47 (vs. 39 Q3 2023)',
    ]
    for h in highlights:
        doc.add_paragraph(h, style='List Bullet')

    # --- Financial Performance ---
    doc.add_heading('2. Financial Performance', level=1)

    doc.add_heading('2.1 Revenue Breakdown', level=2)

    rev_intro = doc.add_paragraph(
        'Revenue composition shifted favorably toward higher-margin recurring streams during the quarter:'
    )

    # Table: Revenue by Segment
    rev_table = doc.add_table(rows=1, cols=4)
    rev_table.style = 'Table Grid'
    hdr = rev_table.rows[0].cells
    for i, h in enumerate(['Segment', 'Q3 2024 ($M)', 'Q3 2023 ($M)', 'YoY Growth']):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    seg_data = [
        ('Enterprise SaaS',   '74.1',  '62.3',  '+19.0%'),
        ('SMB Subscriptions', '31.5',  '28.7',  '+9.8%'),
        ('Professional Svcs', '14.8',  '13.9',  '+6.5%'),
        ('Licensing & Other', '8.0',   '7.4',   '+8.1%'),
        ('Total',             '128.4', '112.3', '+14.2%'),
    ]
    for row_data in seg_data:
        row = rev_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()

    doc.add_heading('2.2 Cost Structure', level=2)
    cost_para = (
        'Total operating expenses were $96.8 million in Q3 2024, compared to $86.9 million in Q3 2023 (+11.4%). '
        'Cost of goods sold (COGS) decreased to 38.7% of revenue (down from 40.2%), reflecting the growing mix '
        'of high-margin SaaS revenue. Sales & marketing spend was $28.4 million (22.1% of revenue), while R&D '
        'investment grew 21.3% to $19.6 million, underscoring commitment to product innovation.'
    )
    doc.add_paragraph(cost_para)

    # --- Operational Metrics ---
    doc.add_heading('3. Operational Metrics', level=1)

    doc.add_heading('3.1 Customer Acquisition & Retention', level=2)
    op_metrics = [
        ('Total Customers',         '4,823',  '4,291',  '+12.4%'),
        ('Enterprise Accounts',     '612',    '547',    '+11.9%'),
        ('Avg. Contract Value ($K)', '121.2', '113.7',  '+6.6%'),
        ('Churn Rate',              '5.8%',   '6.9%',   '-1.1 pp'),
        ('NPS Score',               '67',     '61',     '+6 pts'),
    ]
    op_table = doc.add_table(rows=1, cols=4)
    op_table.style = 'Table Grid'
    op_hdr = op_table.rows[0].cells
    for i, h in enumerate(['Metric', 'Q3 2024', 'Q3 2023', 'Change']):
        op_hdr[i].text = h
        for run in op_hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in op_metrics:
        row = op_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()

    doc.add_heading('3.2 Product Delivery', level=2)
    doc.add_paragraph(
        'The engineering team shipped 3 major product releases and 14 minor updates during Q3 2024. '
        'Platform uptime averaged 99.97%, exceeding the 99.9% SLA commitment. Mean time to resolution (MTTR) '
        'for critical incidents improved to 38 minutes (down from 67 minutes in Q3 2023), attributable to '
        'investments in observability tooling and on-call process improvements.'
    )

    # --- Regional Performance ---
    doc.add_heading('4. Regional Performance', level=1)
    doc.add_paragraph(
        'North America remained the dominant revenue contributor at $81.3 million (63.3% of total revenue). '
        'The EMEA region posted the fastest growth at +22.4% YoY, reaching $28.6 million, driven by enterprise '
        'wins in the UK, Germany, and the Nordics. APAC revenue grew 11.8% to $18.5 million, supported by '
        'expanded channel partnerships in Japan and Australia.'
    )

    reg_table = doc.add_table(rows=1, cols=4)
    reg_table.style = 'Table Grid'
    reg_hdr = reg_table.rows[0].cells
    for i, h in enumerate(['Region', 'Q3 2024 ($M)', 'Q3 2023 ($M)', 'YoY Growth']):
        reg_hdr[i].text = h
        for run in reg_hdr[i].paragraphs[0].runs:
            run.bold = True
    reg_data = [
        ('North America', '81.3', '73.9', '+10.0%'),
        ('EMEA',          '28.6', '23.4', '+22.4%'),
        ('APAC',          '18.5', '16.5', '+11.8%'),
        ('LatAm',         '0.0',  '—',    'N/A'),
    ]
    for row_data in reg_data:
        row = reg_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()

    # --- Outlook ---
    doc.add_heading('5. Q4 2024 Outlook & Strategic Priorities', level=1)
    doc.add_paragraph(
        'Management reaffirms full-year 2024 revenue guidance of $490–$498 million, implying Q4 revenue of '
        '$128–$136 million (+12–19% YoY). Operating margin is expected to expand by approximately 50 basis '
        'points for the full year to 23.8–24.3%. Key strategic initiatives for Q4 include:'
    )
    priorities = [
        'Launch of Acme AI Assistant (GA release scheduled November 12, 2024)',
        'Expansion of the partner channel in EMEA with 12 new GSI agreements',
        'Completion of ISO 27001 certification audit',
        'Integration of DataSync acquisition (closed August 2024) into core platform',
        'Rollout of self-serve onboarding reducing implementation time by 40%',
    ]
    for p in priorities:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        'This document contains forward-looking statements subject to risks and uncertainties. '
        'Actual results may differ materially from those projected. For investor inquiries contact ir@acmecorp.com.'
    )
    for run in disclaimer.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 4. Save as a temporary .docx, then convert to .odt via LibreOffice headless
    tmp_docx = f'{WORKDIR}/_tmp_quarterly_review.docx'
    doc.save(tmp_docx)
    print(f'Temporary .docx saved: {tmp_docx}')

    # Convert .docx -> .odt using LibreOffice headless
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    result = subprocess.run(
        [
            'libreoffice', '--headless', '--convert-to', 'odt',
            '--outdir', REPORTS_DIR,
            tmp_docx,
        ],
        capture_output=True, text=True, env=env, timeout=60
    )
    print('LO convert stdout:', result.stdout)
    print('LO convert stderr:', result.stderr)

    # The output will be named after the tmp file basename; rename if needed
    expected_odt_name = '_tmp_quarterly_review.odt'
    expected_odt_path = f'{REPORTS_DIR}/{expected_odt_name}'
    if os.path.exists(expected_odt_path):
        if os.path.exists(ODT_FILE):
            os.remove(ODT_FILE)
        os.rename(expected_odt_path, ODT_FILE)
        print(f'Renamed: {expected_odt_path} -> {ODT_FILE}')
    elif os.path.exists(ODT_FILE):
        print(f'ODT already at correct path: {ODT_FILE}')
    else:
        # Fallback: LibreOffice may have placed file elsewhere; search
        import glob as glob_module
        found = glob_module.glob(f'{WORKDIR}/**/*quarterly*', recursive=True)
        print(f'Fallback search result: {found}')
        raise RuntimeError(f'ODT conversion failed. stdout: {result.stdout}  stderr: {result.stderr}')

    # Clean up tmp docx
    if os.path.exists(tmp_docx):
        os.remove(tmp_docx)

    # Confirm no PDF exists (initial state)
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
        print(f'Removed stale PDF: {pdf_path}')

    print(f'Initial ODT created: {ODT_FILE}')
    print(f'PDF absent (correct for initial state): {pdf_path}')

    # 5. Open the ODT file in LibreOffice Writer (GUI-ready)
    launch_gui(f'libreoffice --writer "{ODT_FILE}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
