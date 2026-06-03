"""
Initial Setup: Legal brief with widow/orphan control disabled
Task ID: writer_legal_062
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_062'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def disable_widow_orphan_on_style(doc):
    """Explicitly disable widow/orphan control on the Normal (Default Paragraph) style."""
    for style in doc.styles:
        if style.name == 'Normal':
            pPr = style.element.find(qn('w:pPr'))
            if pPr is None:
                pPr = style.element.makeelement(qn('w:pPr'), {})
                style.element.append(pPr)
            # Remove existing widowControl if present
            for wc in pPr.findall(qn('w:widowControl')):
                pPr.remove(wc)
            # Set widowControl to false (disabled)
            wc_elem = pPr.makeelement(qn('w:widowControl'), {qn('w:val'): '0'})
            pPr.insert(0, wc_elem)
            break


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Disable widow/orphan control on the Normal style
    disable_widow_orphan_on_style(doc)

    # --- Title Page ---
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('IN THE UNITED STATES DISTRICT COURT\nFOR THE SOUTHERN DISTRICT OF NEW YORK')
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph('')

    case_info = doc.add_paragraph()
    case_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = case_info.add_run('MERCER TECHNOLOGIES, INC.,\nPlaintiff,\n\nv.\n\nATLANTIC DYNAMICS CORPORATION,\nDefendant.')
    run.font.size = Pt(12)

    doc.add_paragraph('')

    case_no = doc.add_paragraph()
    case_no.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = case_no.add_run('Case No. 24-CV-03891-RMB')
    run.font.size = Pt(12)
    run.bold = True

    doc.add_paragraph('')

    brief_title = doc.add_paragraph()
    brief_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = brief_title.add_run("PLAINTIFF'S MEMORANDUM OF LAW\nIN SUPPORT OF MOTION FOR SUMMARY JUDGMENT")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_page_break()

    # --- Table of Contents ---
    toc_heading = doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_entries = [
        ('I.', 'PRELIMINARY STATEMENT', '1'),
        ('II.', 'STATEMENT OF FACTS', '3'),
        ('III.', 'LEGAL STANDARD', '8'),
        ('IV.', 'ARGUMENT', '10'),
        ('', 'A. The Undisputed Facts Establish Breach of Contract', '10'),
        ('', 'B. Atlantic Dynamics Failed to Perform Under Section 4.2', '14'),
        ('', 'C. Damages Are Readily Calculable', '18'),
        ('', 'D. No Genuine Issue of Material Fact Exists', '22'),
        ('V.', 'CONCLUSION', '27'),
    ]
    for num, title_text, page in toc_entries:
        p = doc.add_paragraph()
        if num:
            run = p.add_run(f'{num}\t{title_text}')
            run.bold = True
        else:
            run = p.add_run(f'\t\t{title_text}')
        p.add_run(f'\t{page}')

    doc.add_page_break()

    # --- Table of Authorities ---
    auth_heading = doc.add_heading('TABLE OF AUTHORITIES', level=1)
    authorities = [
        'Anderson v. Liberty Lobby, Inc., 477 U.S. 242 (1986)',
        'Celotex Corp. v. Catrett, 477 U.S. 317 (1986)',
        'Matsushita Elec. Indus. Co. v. Zenith Radio Corp., 475 U.S. 574 (1986)',
        'Scott v. Harris, 550 U.S. 372 (2007)',
        'Adickes v. S.H. Kress & Co., 398 U.S. 144 (1970)',
        'Reeves v. Sanderson Plumbing Prods., Inc., 530 U.S. 133 (2000)',
        'Tolan v. Cotton, 572 U.S. 650 (2014)',
        'N.Y. U.C.C. § 2-301 et seq.',
        'Fed. R. Civ. P. 56(a)',
        'Fed. R. Civ. P. 56(c)',
    ]
    for auth in authorities:
        p = doc.add_paragraph()
        run = p.add_run(auth)
        run.italic = True
        run.font.size = Pt(11)

    doc.add_page_break()

    # --- I. PRELIMINARY STATEMENT ---
    doc.add_heading('I. PRELIMINARY STATEMENT', level=1)

    legal_paragraphs_section1 = [
        "Plaintiff Mercer Technologies, Inc. (\"Mercer\") respectfully submits this memorandum of law in support of its motion for summary judgment against Defendant Atlantic Dynamics Corporation (\"Atlantic Dynamics\"). This action arises from Atlantic Dynamics' material breach of a Master Services Agreement dated March 15, 2022 (the \"Agreement\"), pursuant to which Atlantic Dynamics was obligated to deliver a customized enterprise resource planning system meeting specified performance benchmarks.",

        "The undisputed record establishes that Atlantic Dynamics failed to deliver the contracted system within the agreed-upon timeline, failed to meet the performance specifications set forth in Schedule B of the Agreement, and ultimately abandoned the project after receiving $4.7 million in milestone payments from Mercer. These failures constitute a clear and material breach of the Agreement under New York law.",

        "Summary judgment is appropriate here because the documentary evidence, including the Agreement itself, project correspondence, third-party audit reports, and Atlantic Dynamics' own internal communications obtained in discovery, conclusively demonstrates that Atlantic Dynamics breached multiple material provisions of the Agreement. There is no genuine dispute as to any material fact, and Mercer is entitled to judgment as a matter of law.",

        "As set forth in greater detail below, the record reflects that Atlantic Dynamics consistently missed contractual deadlines, delivered software modules that failed acceptance testing on multiple occasions, and ultimately notified Mercer on January 8, 2024, that it would not complete the remaining deliverables under the Agreement. The damages flowing from this breach are readily calculable based on the milestone payments made, the cost of remediation by a replacement vendor, and the lost business opportunities documented in Mercer's financial records.",
    ]
    for text in legal_paragraphs_section1:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- II. STATEMENT OF FACTS ---
    doc.add_heading('II. STATEMENT OF FACTS', level=1)

    facts_paragraphs = [
        "Mercer Technologies is a Delaware corporation with its principal place of business in New York, New York. Mercer operates a diversified technology services business with approximately 2,300 employees across fourteen offices in the United States and Canada. In fiscal year 2023, Mercer reported revenues of approximately $890 million. Mercer's core business involves providing managed IT services, cloud infrastructure solutions, and custom software development to enterprise clients in the financial services, healthcare, and manufacturing sectors.",

        "Atlantic Dynamics Corporation is a Virginia corporation with its principal place of business in Arlington, Virginia. Atlantic Dynamics markets itself as a provider of enterprise software solutions, including enterprise resource planning (\"ERP\") systems, customer relationship management platforms, and business intelligence tools. According to its website and marketing materials, Atlantic Dynamics has delivered over 200 enterprise software implementations since its founding in 2015.",

        "In late 2021, Mercer initiated a procurement process to replace its legacy ERP system, which had been in operation since 2014 and was approaching end-of-life. Mercer issued a request for proposals (\"RFP\") to twelve qualified vendors on November 1, 2021. The RFP specified that the replacement system must integrate with Mercer's existing financial reporting infrastructure, support real-time inventory management across all fourteen office locations, and comply with SOC 2 Type II security requirements.",

        "Atlantic Dynamics submitted its proposal on December 15, 2021, representing that it could deliver a fully customized ERP solution meeting all RFP specifications within eighteen months of contract execution, at a total cost of $8.2 million. The proposal included a detailed project plan with six major milestones, each tied to specific deliverables and acceptance criteria. Atlantic Dynamics further represented that it had successfully completed similar implementations for three comparable organizations, including Westfield Financial Group, Cascade Manufacturing Partners, and NorthBridge Health Systems.",

        "Following evaluation of all proposals and multiple rounds of due diligence, including reference checks and technical demonstrations, Mercer selected Atlantic Dynamics as its preferred vendor. On March 15, 2022, the parties executed the Master Services Agreement. The Agreement incorporated Atlantic Dynamics' proposal by reference and established the following key terms:",

        "Section 2.1 of the Agreement required Atlantic Dynamics to deliver the ERP system in six phases, with each phase consisting of defined modules and functionalities. Section 2.3 established specific performance benchmarks for each module, including response time, throughput, and error rate requirements. Section 3.1 set forth the project timeline, with a final delivery date of September 15, 2023. Section 3.2 provided for liquidated damages of $50,000 per week for delays beyond the final delivery date, capped at $2.6 million.",

        "Section 4.1 of the Agreement established the payment schedule, tying payments to the completion of each milestone. Section 4.2 required Atlantic Dynamics to provide written certification upon completion of each milestone, confirming that the deliverables met the acceptance criteria specified in Schedule B. Section 5.1 contained Atlantic Dynamics' representation and warranty that it possessed the technical expertise, personnel, and resources necessary to perform the work described in the Agreement.",

        "The first milestone, covering system architecture design and project planning, was completed on schedule on May 15, 2022. Mercer paid the first milestone payment of $820,000 upon Atlantic Dynamics' written certification that the deliverables met the acceptance criteria. The second milestone, covering database design and core module development, was originally due on August 15, 2022.",

        "On August 1, 2022, Atlantic Dynamics' project manager, David Chen, notified Mercer's project lead, Rachel Morrison, that the second milestone would be delayed by approximately four weeks due to what Atlantic Dynamics characterized as \"unforeseen technical complexity in the data migration architecture.\" Mercer agreed to a two-week extension, and the second milestone was ultimately delivered on September 12, 2022, nearly four weeks late.",

        "Mercer conducted acceptance testing of the second milestone deliverables between September 12 and October 10, 2022. The testing revealed seventeen defects, including three critical defects related to data integrity in the financial reporting module. Mercer notified Atlantic Dynamics of the defects on October 12, 2022, and Atlantic Dynamics represented that all defects would be resolved within thirty days. The defects were not fully resolved until December 20, 2022, more than two months after notification.",

        "The third milestone, covering integration module development and API connectivity, was originally due on November 15, 2022. Due to the cascading delays from the second milestone, Atlantic Dynamics requested a revised timeline on November 1, 2022. The parties negotiated an amendment to the Agreement (\"Amendment No. 1\"), executed on November 20, 2022, which extended the third milestone deadline to February 15, 2023, and the final delivery date to December 15, 2023. Amendment No. 1 did not modify the liquidated damages provision.",

        "Atlantic Dynamics delivered the third milestone on March 8, 2023, three weeks past the revised deadline. Acceptance testing revealed twenty-three defects, including five critical defects. Among the critical defects were failures in the real-time inventory synchronization module, which produced inconsistent data across locations when processing concurrent transactions. Atlantic Dynamics acknowledged the defects and committed to resolution by April 30, 2023. The critical defects were not fully resolved until June 15, 2023.",

        "The fourth milestone, covering user interface development and reporting dashboards, was originally due on April 15, 2023, under the revised timeline. Atlantic Dynamics delivered the fourth milestone on July 22, 2023, more than three months late. The delivered modules exhibited significant performance issues, with response times exceeding the contractual benchmarks by 300% to 500%. Mercer's independent technical consultant, Dr. Patricia Langford of Langford Technology Advisors, reviewed the delivered code and concluded that the performance issues were attributable to fundamental architectural decisions that would require substantial refactoring to remediate.",

        "On August 15, 2023, Mercer sent Atlantic Dynamics a formal notice of default under Section 7.2 of the Agreement, citing the persistent delays, the accumulation of unresolved defects, and the failure to meet performance benchmarks. The notice provided Atlantic Dynamics with sixty days to cure the defaults, as required by Section 7.3 of the Agreement.",

        "Atlantic Dynamics responded on August 28, 2023, acknowledging the delays but disputing the characterization of its performance as a material breach. Atlantic Dynamics proposed a revised project plan that would extend the final delivery date to June 30, 2024, and offered a $500,000 credit against future milestone payments as consideration for Mercer's agreement to the extension. Mercer rejected this proposal on September 10, 2023, noting that the proposed extension would result in a total project duration more than double the original estimate.",

        "Between September and December 2023, Atlantic Dynamics continued work on the project but made limited progress. The fifth milestone, covering system integration testing and security compliance verification, remained substantially incomplete. Atlantic Dynamics' project team was reduced from fourteen full-time engineers to six during this period, as several key personnel were reassigned to other client engagements.",

        "On January 8, 2024, Atlantic Dynamics' Chief Technology Officer, Margaret Sullivan, sent a letter to Mercer's General Counsel, James Whitaker, stating that Atlantic Dynamics was \"unable to complete the remaining deliverables under the Agreement within a commercially reasonable timeframe\" and proposing a \"mutual termination\" of the Agreement with each party bearing its own costs. This communication constituted an anticipatory repudiation of the Agreement under New York law.",

        "As of January 8, 2024, Mercer had paid Atlantic Dynamics a total of $4,756,000 in milestone payments, representing payments for the first four milestones. Mercer subsequently engaged Pinnacle Systems Group (\"Pinnacle\") to assess the state of the partially completed system and develop a remediation plan. Pinnacle's assessment, completed on March 15, 2024, concluded that approximately 35% of Atlantic Dynamics' delivered code was salvageable, that remediation and completion of the system would require an additional $6.8 million, and that the projected completion date was December 31, 2024.",

        "Mercer also retained FairValue Economics LLC to quantify the damages resulting from Atlantic Dynamics' breach. FairValue's expert report, dated April 30, 2024, calculated Mercer's total damages at $14.2 million, comprising: (a) $4,756,000 in milestone payments to Atlantic Dynamics; (b) $6,800,000 in remediation costs payable to Pinnacle; and (c) $2,644,000 in lost business opportunities attributable to the delayed ERP implementation, based on analysis of contracts that Mercer was unable to bid on or fulfill due to capacity constraints arising from the failed implementation.",
    ]
    for text in facts_paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- III. LEGAL STANDARD ---
    doc.add_heading('III. LEGAL STANDARD', level=1)

    standard_paragraphs = [
        "Summary judgment is appropriate when \"the movant shows that there is no genuine dispute as to any material fact and the movant is entitled to judgment as a matter of law.\" Fed. R. Civ. P. 56(a). The movant bears the initial burden of demonstrating the absence of a genuine issue of material fact. Celotex Corp. v. Catrett, 477 U.S. 317, 323 (1986). Once the movant has met this burden, the nonmoving party must set forth specific facts showing that there is a genuine issue for trial. Anderson v. Liberty Lobby, Inc., 477 U.S. 242, 248 (1986).",

        "A fact is \"material\" if it \"might affect the outcome of the suit under the governing law.\" Anderson, 477 U.S. at 248. A dispute about a material fact is \"genuine\" if \"the evidence is such that a reasonable jury could return a verdict for the nonmoving party.\" Id. In determining whether a genuine issue of material fact exists, the court must view the evidence in the light most favorable to the nonmoving party and draw all reasonable inferences in its favor. Matsushita Elec. Indus. Co. v. Zenith Radio Corp., 475 U.S. 574, 587 (1986).",

        "However, the nonmoving party \"must do more than simply show that there is some metaphysical doubt as to the material facts.\" Matsushita, 475 U.S. at 586. The mere existence of a scintilla of evidence in support of the nonmoving party's position is insufficient to survive summary judgment. Anderson, 477 U.S. at 252. Rather, the nonmoving party must present \"significant probative evidence\" demonstrating that a genuine factual dispute exists. Id. at 249.",

        "Under New York law, a breach of contract claim requires proof of four elements: (1) the existence of a valid contract; (2) performance by the plaintiff; (3) breach by the defendant; and (4) damages resulting from the breach. Harris v. Seward Park Hous. Corp., 79 A.D.3d 425, 426 (1st Dep't 2010). Where the contract terms are clear and unambiguous, the construction of the contract is a matter of law for the court. W.W.W. Assocs., Inc. v. Giancontieri, 77 N.Y.2d 157, 162 (1990).",

        "An anticipatory repudiation occurs when a party to a contract makes a \"definite and unequivocal\" manifestation of intent not to perform its obligations under the contract. Norcon Power Partners, L.P. v. Niagara Mohawk Power Corp., 92 N.Y.2d 458, 463 (1998). When a party anticipatorily repudiates a contract, the non-breaching party may treat the repudiation as an immediate breach and pursue remedies accordingly. Id.",
    ]
    for text in standard_paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- IV. ARGUMENT ---
    doc.add_heading('IV. ARGUMENT', level=1)

    doc.add_heading('A. The Undisputed Facts Establish Breach of Contract', level=2)

    argument_a = [
        "The four elements of breach of contract are established as a matter of law on this record. First, the existence of the Agreement is undisputed. Both parties have acknowledged the validity and enforceability of the Master Services Agreement dated March 15, 2022, and Amendment No. 1 dated November 20, 2022. The Agreement is a fully integrated contract that clearly sets forth the parties' rights and obligations.",

        "Second, Mercer fully performed its obligations under the Agreement. Mercer timely made all milestone payments totaling $4,756,000 upon Atlantic Dynamics' certification that each milestone had been completed. Mercer provided all required access to its systems, facilities, and personnel as specified in Section 6.1 of the Agreement. Mercer also complied with its obligations to provide timely feedback during acceptance testing and to negotiate in good faith regarding schedule modifications.",

        "Third, Atlantic Dynamics materially breached the Agreement in multiple respects. The undisputed evidence demonstrates that Atlantic Dynamics: (a) consistently failed to meet contractual deadlines for milestone deliveries; (b) delivered software modules that failed to meet the performance benchmarks specified in Schedule B; (c) failed to provide adequate staffing and resources to complete the project, in breach of its representations under Section 5.1; and (d) ultimately repudiated the Agreement through its January 8, 2024 letter.",

        "Fourth, Mercer has suffered quantifiable damages as a direct result of Atlantic Dynamics' breach, including the milestone payments made, remediation costs, and lost business opportunities, totaling $14.2 million as documented in the FairValue expert report.",

        "Each of these elements is supported by documentary evidence that Atlantic Dynamics cannot genuinely dispute. The Agreement speaks for itself. The milestone payment records are contemporaneous business records. The acceptance testing results were generated jointly by both parties' technical teams. And Atlantic Dynamics' own January 8, 2024 letter constitutes an unequivocal admission that it cannot perform.",
    ]
    for text in argument_a:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('B. Atlantic Dynamics Failed to Perform Under Section 4.2', level=2)

    argument_b = [
        "Section 4.2 of the Agreement required Atlantic Dynamics to provide written certification upon completion of each milestone, confirming that the deliverables met the acceptance criteria specified in Schedule B. This provision was a material term of the Agreement because it served as the trigger for Mercer's payment obligations and as a quality assurance mechanism.",

        "The record demonstrates that Atlantic Dynamics' certifications for milestones two through four were materially inaccurate. For the second milestone, Atlantic Dynamics certified completion on September 12, 2022, but acceptance testing revealed seventeen defects, including three critical defects. For the third milestone, Atlantic Dynamics certified completion on March 8, 2023, but acceptance testing revealed twenty-three defects, including five critical defects. For the fourth milestone, Atlantic Dynamics certified completion on July 22, 2023, but the delivered modules exhibited response times 300% to 500% above contractual benchmarks.",

        "These false certifications constitute independent breaches of the Agreement. Under New York law, a party's misrepresentation of its performance under a contract may give rise to claims for both breach of contract and fraud. Merrill Lynch & Co. v. Allegheny Energy, Inc., 500 F.3d 171, 183 (2d Cir. 2007). While Mercer does not assert a fraud claim in this action, the pattern of false certifications is relevant to establishing the materiality of Atlantic Dynamics' breach and to rebutting any defense that the defects were minor or inconsequential.",

        "The deposition testimony of Atlantic Dynamics' project manager, David Chen, further confirms the knowing nature of these false certifications. Mr. Chen testified that he was aware of multiple unresolved defects at the time he signed the second milestone certification, but that he was instructed by his supervisor, Atlantic Dynamics' Vice President of Engineering, Robert Hayes, to issue the certification in order to trigger the milestone payment. (Chen Dep. Tr. at 147:3-22.) This testimony is devastating to any defense that Atlantic Dynamics acted in good faith.",

        "Moreover, the internal emails produced in discovery reveal that Atlantic Dynamics' leadership was aware as early as October 2022 that the project was fundamentally understaffed and that the contractual timeline was unachievable. A November 3, 2022 email from CTO Margaret Sullivan to CEO Thomas Bradley states: \"We are significantly behind on the Mercer project and I do not see a realistic path to meeting the revised timeline without at least eight additional senior engineers, which we do not have available.\" (Ex. 14.) Despite this awareness, Atlantic Dynamics continued to accept milestone payments and represent that the project was on track.",

        "Atlantic Dynamics may attempt to argue that Mercer's acceptance of the milestone deliverables, notwithstanding the identified defects, constituted a waiver of the performance standards. This argument fails for two reasons. First, Section 9.4 of the Agreement contains an express anti-waiver provision stating that \"no waiver of any provision of this Agreement shall be effective unless in writing and signed by the party against whom it is sought to be enforced.\" Mercer never executed any written waiver of the acceptance criteria.",

        "Second, even absent the anti-waiver clause, Mercer's conditional acceptance of deficient deliverables, accompanied by contemporaneous written objections and cure demands, does not constitute a waiver under New York law. A waiver requires a voluntary and intentional relinquishment of a known right. Gilbert Frank Corp. v. Federal Ins. Co., 70 N.Y.2d 966, 968 (1988). Mercer's consistent pattern of identifying defects, demanding cure, and reserving its rights is the antithesis of waiver.",
    ]
    for text in argument_b:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('C. Damages Are Readily Calculable', level=2)

    argument_c = [
        "Mercer's damages are supported by comprehensive documentary evidence and expert analysis. The FairValue Economics report quantifies Mercer's total damages at $14.2 million, comprising three categories: milestone payments, remediation costs, and lost business opportunities.",

        "The first category, milestone payments of $4,756,000, is established by bank records and payment confirmations that are beyond genuine dispute. These payments were made in reliance on Atlantic Dynamics' certifications and are recoverable as restitution under New York law. Clark-Fitzpatrick, Inc. v. Long Island R.R. Co., 70 N.Y.2d 382, 389 (1987).",

        "The second category, remediation costs of $6,800,000, is documented in Pinnacle Systems Group's assessment report and corresponding engagement agreement. Pinnacle's assessment was conducted by a team of twelve engineers over six weeks and provides a detailed, line-item budget for completing and remediating the ERP system. Atlantic Dynamics has not challenged the qualifications of Pinnacle's team or the methodology of their assessment.",

        "The third category, lost business opportunities of $2,644,000, is supported by FairValue Economics' analysis of Mercer's bid history, capacity utilization data, and financial projections. FairValue's expert, Dr. Sandra Whitmore, holds a Ph.D. in Economics from Columbia University and has over twenty years of experience in commercial damages analysis. Dr. Whitmore analyzed twelve specific contracts that Mercer declined to bid on or was unable to fulfill during the period of the failed implementation, and calculated the expected profit margin for each based on Mercer's historical performance on comparable engagements.",

        "Under New York law, consequential damages for breach of contract are recoverable where they were reasonably foreseeable at the time of contracting. Kenford Co. v. County of Erie, 73 N.Y.2d 312, 319 (1989). Here, it was entirely foreseeable that a failed ERP implementation would impair Mercer's operational capacity and result in lost business opportunities. The Agreement itself acknowledges the critical nature of the system to Mercer's operations, stating in the recitals that \"Mercer requires a modern ERP system to support its continued growth and operational efficiency.\"",

        "In addition to compensatory damages, Mercer is entitled to recover the liquidated damages specified in Section 3.2 of the Agreement. The final delivery date under the revised timeline was December 15, 2023. Atlantic Dynamics' January 8, 2024 repudiation establishes that performance did not occur by that date, triggering the liquidated damages provision. At $50,000 per week, the liquidated damages from December 15, 2023, through January 8, 2024, total $175,000. When added to the compensatory damages, Mercer's total recoverable damages are $14,375,000.",
    ]
    for text in argument_c:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('D. No Genuine Issue of Material Fact Exists', level=2)

    argument_d = [
        "Atlantic Dynamics cannot identify any genuine issue of material fact that would preclude summary judgment. The existence and terms of the Agreement are undisputed. The timeline of milestone deliveries and the results of acceptance testing are documented in contemporaneous records that Atlantic Dynamics has not challenged. Atlantic Dynamics' January 8, 2024 letter is an unambiguous admission that it cannot perform. And Mercer's damages are supported by expert analysis based on verifiable financial data.",

        "To the extent Atlantic Dynamics may argue that it was excused from performance by Mercer's conduct, such arguments are refuted by the record. Atlantic Dynamics has not identified any specific obligation that Mercer failed to perform under the Agreement. To the contrary, the contemporaneous correspondence reflects Mercer's consistent cooperation and willingness to accommodate reasonable schedule adjustments.",

        "Atlantic Dynamics may also attempt to raise a defense based on the doctrine of commercial impracticability. However, commercial impracticability requires an unforeseen supervening event that makes performance impracticable. Kel Kim Corp. v. Central Markets, Inc., 70 N.Y.2d 900, 902 (1987). Here, the difficulties Atlantic Dynamics encountered were not unforeseen supervening events but rather the natural consequences of inadequate staffing, poor project management, and technical deficiencies in Atlantic Dynamics' own approach. These are precisely the types of business risks that Atlantic Dynamics assumed when it executed the Agreement.",

        "Furthermore, Atlantic Dynamics' reliance on any impracticability defense is undermined by the evidence that its difficulties were foreseeable and, indeed, foreseen. As noted above, CTO Sullivan's November 2022 email acknowledged that the project was understaffed and that the timeline was unrealistic. Yet Atlantic Dynamics chose to continue accepting payments rather than disclosing these issues to Mercer and negotiating a genuine resolution.",

        "The doctrine of substantial performance is similarly unavailable to Atlantic Dynamics. Substantial performance applies only where a party has made a good-faith effort to perform and the deviations from the contract specifications are minor. Jacob & Youngs, Inc. v. Kent, 230 N.Y. 239, 244 (1921). Here, the deviations from the contract specifications were far from minor. Atlantic Dynamics failed to deliver two of six milestones entirely, and the milestones it did deliver contained critical defects and severe performance deficiencies. This is not a case of minor imperfections in an otherwise completed project; it is a case of fundamental failure to perform.",

        "Finally, Atlantic Dynamics cannot create a genuine issue of material fact through conclusory allegations or speculation. \"The mere existence of a scintilla of evidence in support of the plaintiff's position will be insufficient; there must be evidence on which the jury could reasonably find for the plaintiff.\" Anderson, 477 U.S. at 252. Atlantic Dynamics has had the benefit of full discovery, including document production, interrogatories, and depositions of Mercer's key personnel. It has failed to develop any evidence that would support a defense to Mercer's breach of contract claim.",

        "In its discovery responses and deposition testimony, Atlantic Dynamics has not identified any contractual provision that Mercer breached, any factual basis for excusing Atlantic Dynamics' non-performance, or any deficiency in Mercer's damages calculation. Instead, Atlantic Dynamics has relied on vague assertions that the project was \"more complex than anticipated\" and that Mercer's \"evolving requirements\" contributed to the delays. These assertions are both unsupported and irrelevant. The Agreement contains a fixed scope of work, and Mercer has never requested any modifications outside the formal amendment process.",
    ]
    for text in argument_d:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- V. CONCLUSION ---
    doc.add_heading('V. CONCLUSION', level=1)

    conclusion_paragraphs = [
        "For the foregoing reasons, Mercer Technologies respectfully requests that the Court grant its motion for summary judgment on its breach of contract claim, enter judgment in favor of Mercer and against Atlantic Dynamics in the amount of $14,375,000, plus pre-judgment interest at the statutory rate from January 8, 2024, and award Mercer its costs and attorneys' fees incurred in this action as permitted by Section 8.3 of the Agreement.",

        "The undisputed record demonstrates that Atlantic Dynamics materially breached the Master Services Agreement through persistent delays, delivery of deficient work product, false milestone certifications, and ultimate repudiation of its contractual obligations. Mercer performed all of its obligations under the Agreement and has suffered substantial, documented damages as a direct result of Atlantic Dynamics' breach. No genuine issue of material fact exists, and Mercer is entitled to judgment as a matter of law.",
    ]
    for text in conclusion_paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # Signature block
    doc.add_paragraph('')
    sig = doc.add_paragraph('Respectfully submitted,')
    sig.paragraph_format.space_after = Pt(24)

    firm = doc.add_paragraph()
    run = firm.add_run('HARRISON, WHITAKER & COLE LLP')
    run.bold = True

    doc.add_paragraph('')

    attorney = doc.add_paragraph('By: ________________________________')
    doc.add_paragraph('James A. Whitaker, Esq.')
    doc.add_paragraph('Sarah L. Morrison, Esq.')
    doc.add_paragraph('Harrison, Whitaker & Cole LLP')
    doc.add_paragraph('450 Lexington Avenue, 38th Floor')
    doc.add_paragraph('New York, New York 10017')
    doc.add_paragraph('Telephone: (212) 555-7890')
    doc.add_paragraph('Facsimile: (212) 555-7891')
    doc.add_paragraph('Email: jwhitaker@hwclaw.com')

    doc.add_paragraph('')
    doc.add_paragraph('Attorneys for Plaintiff Mercer Technologies, Inc.')

    # Ensure all paragraphs have widow/orphan control disabled
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = para._element.makeelement(qn('w:pPr'), {})
            para._element.insert(0, pPr)
        # Remove existing widowControl
        for wc in pPr.findall(qn('w:widowControl')):
            pPr.remove(wc)
        # Set to disabled
        wc = pPr.makeelement(qn('w:widowControl'), {qn('w:val'): '0'})
        pPr.insert(0, wc)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
