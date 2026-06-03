"""
Reward Script: Apply widow and orphan control to body paragraphs in a legal brief
Task ID: writer_legal_062
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Normal style has widowControl enabled
  Component 2 (0.4): Majority of paragraphs have effective widow control (True or inherit-from-enabled-style)
  Component 3 (0.2): No paragraphs explicitly have widow control disabled
"""

import os
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_062'
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_normal_style_widow_control(doc):
    """
    Check if the Normal (default paragraph) style has widowControl enabled.
    In OOXML:
      <w:widowControl/>          -> enabled (no val = true)
      <w:widowControl w:val="1"/> -> enabled
      <w:widowControl w:val="0"/> -> disabled
      absent                     -> default is enabled per spec, but we check explicitly
    Returns: True if enabled, False if disabled or absent
    """
    for style_el in doc.styles.element.findall('.//{%s}style' % WNS):
        style_id = style_el.attrib.get('{%s}styleId' % WNS, '')
        if style_id == 'Normal':
            ppr = style_el.find('{%s}pPr' % WNS)
            if ppr is not None:
                wc = ppr.find('{%s}widowControl' % WNS)
                if wc is not None:
                    val = wc.attrib.get('{%s}val' % WNS, None)
                    # No val attribute or val != "0" means enabled
                    if val is None or val != '0':
                        return True
                    else:
                        return False
                else:
                    # widowControl not present in style pPr
                    # Per OOXML spec default is enabled, but for scoring
                    # we want explicit presence
                    return False
            else:
                return False
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Normal style has widowControl enabled (0.4 points)
    # This should FAIL on initial (val="0") and PASS on golden (no val = enabled)
    try:
        style_wc_enabled = get_normal_style_widow_control(doc)
        if style_wc_enabled:
            print(f"PASS: Component 1 -- Normal style widowControl is enabled (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 -- Normal style widowControl is NOT enabled")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Majority of body paragraphs have effective widow control (0.4 points)
    # "Effective" means either explicitly True, or None (inheriting from the enabled style)
    # On initial: all 116 are False -> 0% effective -> FAIL
    # On golden: 71 True + 45 None (inheriting enabled) -> 100% effective -> PASS
    try:
        total_paras = len(doc.paragraphs)
        effective_wc_count = 0
        for p in doc.paragraphs:
            wc = p.paragraph_format.widow_control
            if wc is True:
                effective_wc_count += 1
            elif wc is None and style_wc_enabled:
                # Inherits from style, which is enabled
                effective_wc_count += 1

        if total_paras > 0:
            ratio = effective_wc_count / total_paras
            print(f"  Effective widow control: {effective_wc_count}/{total_paras} = {ratio:.2%}")
            if ratio >= 0.9:
                print(f"PASS: Component 2 -- >=90% of paragraphs have effective widow control (0.4 pts)")
                total_score += 0.4
            elif ratio >= 0.5:
                partial = 0.4 * (ratio - 0.5) / 0.4  # scale 0.5-0.9 to 0-1
                print(f"PARTIAL: Component 2 -- {ratio:.0%} paragraphs have effective widow control ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 -- Only {ratio:.0%} of paragraphs have effective widow control")
        else:
            print(f"FAIL: Component 2 -- No paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: No paragraphs explicitly have widow control disabled (0.2 points)
    # On initial: all 116 are explicitly False -> FAIL
    # On golden: 0 are explicitly False -> PASS
    try:
        disabled_count = 0
        for p in doc.paragraphs:
            wc = p.paragraph_format.widow_control
            if wc is False:
                disabled_count += 1

        if disabled_count == 0:
            print(f"PASS: Component 3 -- No paragraphs have widow control explicitly disabled (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 -- {disabled_count} paragraphs have widow control explicitly disabled")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
