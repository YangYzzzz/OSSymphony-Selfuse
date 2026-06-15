"""
Initial Setup: Constitutional Law Memorandum without footnotes
Task ID: writer_legal_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_051'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('MEMORANDUM OF LAW', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Case caption block
    caption = doc.add_paragraph()
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = caption.add_run('IN THE MATTER OF:')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    caption2 = doc.add_paragraph()
    run2 = caption2.add_run('State of Columbia v. Henderson Municipal School District')
    run2.italic = True
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    caption3 = doc.add_paragraph()
    run3 = caption3.add_run('Case No. 2025-CV-04817')
    run3.font.size = Pt(12)
    run3.font.name = 'Times New Roman'

    # Horizontal line (border)
    doc.add_paragraph('_' * 72)

    # Section I
    h1 = doc.add_heading('I. INTRODUCTION', level=1)

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1.5
    r = p1.add_run(
        'This memorandum is submitted on behalf of Petitioners in the above-captioned '
        'matter. The central question before this Court concerns the extent to which '
        'the Henderson Municipal School District\'s revised admissions policy infringes '
        'upon the due process rights of applicants and their families under the '
        'Fourteenth Amendment to the United States Constitution.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    r = p2.add_run(
        'The policy at issue categorically excludes certain applicants based on '
        'residential zoning classifications without providing any individualized '
        'assessment or meaningful opportunity to be heard. As set forth below, '
        'this practice violates the guarantee of equal protection enshrined in '
        'the Fourteenth Amendment and reinforced by decades of binding precedent.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Section II
    doc.add_heading('II. STATEMENT OF FACTS', level=1)

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1.5
    r = p3.add_run(
        'In March 2025, the Henderson Municipal School District Board of Trustees '
        'adopted Resolution No. 2025-12, which established a tiered admissions system '
        'for its magnet school program. Under the new policy, applicants residing in '
        'Zone A receive automatic admission, while applicants from Zones B and C must '
        'satisfy additional criteria including a standardized assessment score above '
        'the 85th percentile and a demonstrated record of extracurricular achievement.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    p4 = doc.add_paragraph()
    p4.paragraph_format.line_spacing = 1.5
    r = p4.add_run(
        'Petitioners Reyes, Okafor, and Tanaka are families residing in Zone C whose '
        'children were denied admission to Westbrook Magnet Academy despite strong '
        'academic records. The families contend that the zone-based tier system '
        'operates as a proxy for socioeconomic status and disproportionately impacts '
        'minority communities concentrated in Zone C.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Section III
    doc.add_heading('III. ARGUMENT', level=1)

    doc.add_heading('A. The Revised Policy Fails Strict Scrutiny', level=2)

    p5 = doc.add_paragraph()
    p5.paragraph_format.line_spacing = 1.5
    r = p5.add_run(
        'When a government classification burdens a fundamental right or targets a '
        'suspect class, courts apply the standard of strict scrutiny to determine '
        'whether the classification is constitutionally permissible. Under this exacting '
        'standard, the government must demonstrate that the challenged policy serves a '
        'compelling interest and is narrowly tailored to achieve that interest.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    p6 = doc.add_paragraph()
    p6.paragraph_format.line_spacing = 1.5
    r = p6.add_run(
        'The District has asserted that the tiered admissions system promotes '
        '"neighborhood cohesion" and "efficient resource allocation." However, '
        'neither of these stated objectives rises to the level of a compelling interest '
        'sufficient to justify a classification scheme that effectively segregates '
        'students along socioeconomic and racial lines.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    doc.add_heading('B. The Policy Also Fails Rational Basis Review', level=2)

    p7 = doc.add_paragraph()
    p7.paragraph_format.line_spacing = 1.5
    r = p7.add_run(
        'Even under the more deferential standard of rational basis review, the '
        'District\'s policy cannot withstand scrutiny. A classification satisfies '
        'rational basis only if it bears a reasonable relationship to a legitimate '
        'governmental purpose. Here, the purported connection between residential '
        'zoning and educational merit is tenuous at best.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    p8 = doc.add_paragraph()
    p8.paragraph_format.line_spacing = 1.5
    r = p8.add_run(
        'The District has offered no empirical evidence demonstrating that Zone A '
        'residents are more likely to succeed in the magnet program than students from '
        'other zones. In fact, historical performance data shows that Zone C students '
        'who gained admission under the prior policy performed at or above the district '
        'average. The zone-based classification is therefore arbitrary and irrational.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Section IV
    doc.add_heading('IV. CONCLUSION', level=1)

    p9 = doc.add_paragraph()
    p9.paragraph_format.line_spacing = 1.5
    r = p9.add_run(
        'For the foregoing reasons, Petitioners respectfully request that this Court '
        'declare Resolution No. 2025-12 unconstitutional and enjoin the Henderson '
        'Municipal School District from enforcing the tiered admissions policy. The '
        'policy infringes upon the due process rights and equal protection guarantees '
        'afforded to all citizens under the Fourteenth Amendment.'
    )
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Signature block
    doc.add_paragraph('')
    sig = doc.add_paragraph()
    sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = sig.add_run('Respectfully submitted,')
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    doc.add_paragraph('')
    atty = doc.add_paragraph()
    r = atty.add_run('Elena M. Vasquez, Esq.')
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    firm = doc.add_paragraph()
    r = firm.add_run('Vasquez & Chen LLP')
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    addr = doc.add_paragraph()
    r = addr.add_run('1200 Constitution Avenue, Suite 400\nColumbia, ST 60201\n(555) 234-8901')
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
