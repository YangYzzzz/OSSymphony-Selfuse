"""
Initial Setup: Office memo with signature block (no indent on signature)
Task ID: writer_para_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'office_memo'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Paragraph 1: 'MEMORANDUM' (center-aligned, bold)
    p1 = doc.add_paragraph()
    run1 = p1.add_run('MEMORANDUM')
    run1.bold = True
    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1.paragraph_format.left_indent = Cm(0)

    # Paragraph 2
    p2 = doc.add_paragraph('TO: All Department Managers')
    p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p2.paragraph_format.left_indent = Cm(0)

    # Paragraph 3
    p3 = doc.add_paragraph('FROM: Chief Operations Officer')
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p3.paragraph_format.left_indent = Cm(0)

    # Paragraph 4
    p4 = doc.add_paragraph('DATE: March 3, 2025')
    p4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p4.paragraph_format.left_indent = Cm(0)

    # Paragraph 5
    p5 = doc.add_paragraph('RE: Updated Travel Expense Reimbursement Policy')
    p5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p5.paragraph_format.left_indent = Cm(0)

    # Paragraph 6
    p6 = doc.add_paragraph(
        'Effective April 1, 2025, all travel expense reimbursement requests must be submitted '
        'through the new digital portal within 14 business days of travel completion. '
        'Paper forms will no longer be accepted.'
    )
    p6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p6.paragraph_format.left_indent = Cm(0)

    # Paragraph 7
    p7 = doc.add_paragraph(
        'Please ensure your teams are informed of this change and have access to the portal by March 15.'
    )
    p7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p7.paragraph_format.left_indent = Cm(0)

    # Paragraph 8: Signature block - 'Best regards,' (NO indent in initial state)
    p8 = doc.add_paragraph('Best regards,')
    p8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p8.paragraph_format.left_indent = Cm(0)

    # Paragraph 9: 'Patricia Chen' (NO indent in initial state)
    p9 = doc.add_paragraph('Patricia Chen')
    p9.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p9.paragraph_format.left_indent = Cm(0)

    # Paragraph 10: 'Chief Operations Officer' (NO indent in initial state)
    p10 = doc.add_paragraph('Chief Operations Officer')
    p10.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p10.paragraph_format.left_indent = Cm(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the initial file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
