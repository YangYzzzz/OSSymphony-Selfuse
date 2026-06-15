"""
Initial Setup: Create a Writer document with 7 product features as plain text paragraphs
Task ID: writer_lec_015
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("NexaHub Pro Smart Home Controller", level=1)

    # Product description paragraph
    intro = doc.add_paragraph(
        "The NexaHub Pro is our flagship smart home controller, designed to "
        "seamlessly integrate all your connected devices into a single, "
        "intuitive ecosystem. Whether you are managing lighting, climate, "
        "security, or entertainment, the NexaHub Pro puts everything at your "
        "fingertips."
    )
    intro.paragraph_format.space_after = Pt(6)

    # Subheading for features
    doc.add_heading("Key Features", level=2)

    # 7 features as plain text paragraphs (NO bullets, NO arrows, NO list styles)
    features = [
        "Supports over 200 smart home brands including Philips Hue, Nest, Ring, and Sonos with automatic device discovery",
        "Voice control integration with Amazon Alexa, Google Assistant, and Apple Siri for hands-free operation",
        "Advanced energy monitoring dashboard that tracks real-time power consumption across all connected devices",
        "Military-grade AES-256 encryption for all device communications with automatic security patch updates",
        "Customizable automation routines that trigger actions based on time, location, weather, or sensor data",
        "Built-in 7-inch touchscreen display with ambient light sensor for comfortable viewing day and night",
        "Remote access through the NexaHub mobile app allowing full home control from anywhere in the world",
    ]

    for feature_text in features:
        para = doc.add_paragraph(feature_text)
        para.paragraph_format.space_after = Pt(4)

    # Closing paragraph
    closing = doc.add_paragraph(
        "Experience the future of smart living with NexaHub Pro. Available at "
        "authorized retailers and online at nexahub.com starting March 2026."
    )
    closing.paragraph_format.space_before = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
