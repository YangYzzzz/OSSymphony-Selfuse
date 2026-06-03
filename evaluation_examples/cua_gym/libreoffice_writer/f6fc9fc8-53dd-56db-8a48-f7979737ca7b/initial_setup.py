"""
Initial Setup: Legal contract document with continuous page numbering 1-22
Task ID: writer_legal_044
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_044'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section):
    """Add page number field to footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # PAGE field code
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)

    for r in [r1, r2, r3]:
        r.font.size = Pt(10)
        r.font.name = 'Times New Roman'


def add_heading(doc, text, level=1):
    """Add a heading with legal document styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_clause(doc, text, bold_first_sentence=False):
    """Add a paragraph of legal text."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15

    if bold_first_sentence:
        # Split at first period
        idx = text.find('.')
        if idx > 0:
            r1 = para.add_run(text[:idx + 1])
            r1.bold = True
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(11)
            r2 = para.add_run(text[idx + 1:])
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(11)
        else:
            r = para.add_run(text)
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
    else:
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
    return para


def add_filler_text(doc, count=4):
    """Add multiple paragraphs of realistic legal text to fill pages."""
    legal_paragraphs = [
        "The Parties hereby acknowledge and agree that the obligations set forth in this Agreement shall be binding upon and inure to the benefit of each Party and their respective successors, assigns, heirs, executors, administrators, and legal representatives. No assignment of this Agreement or any rights or obligations hereunder may be made by either Party without the prior written consent of the other Party.",
        "Notwithstanding any other provision of this Agreement to the contrary, neither Party shall be liable to the other for any indirect, incidental, consequential, special, or exemplary damages arising out of or related to this Agreement, including but not limited to loss of revenue, loss of profits, loss of business, or loss of data, even if such Party has been advised of the possibility of such damages.",
        "Each Party represents and warrants that: (a) it has full corporate power and authority to execute, deliver, and perform this Agreement; (b) the execution, delivery, and performance of this Agreement has been duly authorized by all necessary corporate action; (c) this Agreement constitutes a legal, valid, and binding obligation enforceable against such Party in accordance with its terms.",
        "In the event of any dispute, claim, question, or disagreement arising from or relating to this Agreement or the breach thereof, the Parties hereto shall use their best efforts to settle the dispute through consultation and negotiation in good faith and a spirit of mutual cooperation. If those efforts fail to resolve the dispute within thirty (30) days after written notice of the dispute is delivered, the dispute shall be submitted to binding arbitration.",
        "This Agreement may be executed in counterparts, each of which shall be deemed an original, but all of which together shall constitute one and the same instrument. Delivery of an executed counterpart of a signature page to this Agreement by facsimile or electronic transmission shall be effective as delivery of a manually executed counterpart.",
        "All notices, requests, consents, claims, demands, waivers, and other communications hereunder shall be in writing and shall be deemed to have been given: (a) when delivered by hand; (b) when received by the addressee if sent by a nationally recognized overnight courier; (c) on the date sent by email if sent during normal business hours of the recipient.",
        "The failure of either Party to enforce any provision of this Agreement shall not constitute a waiver of future enforcement of that or any other provision of this Agreement. No waiver of any provision of this Agreement shall be effective unless it is in the form of a writing signed by the Party granting such waiver.",
        "This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware, without regard to its conflict of laws provisions. The Parties irrevocably submit to the exclusive jurisdiction of the federal and state courts located in Wilmington, Delaware for the resolution of any disputes arising under this Agreement.",
        "If any provision of this Agreement is found to be invalid, illegal, or unenforceable by a court of competent jurisdiction, then such provision shall be modified to the minimum extent necessary to make it valid, legal, and enforceable while preserving the Parties' original intent as closely as possible.",
        "The Receiving Party shall hold and maintain in strict confidence all Confidential Information disclosed by the Disclosing Party for a period of five (5) years from the date of disclosure. The Receiving Party shall not, without the prior written approval of the Disclosing Party, use for any purpose or disclose to any third party any Confidential Information.",
        "Upon termination or expiration of this Agreement for any reason, each Party shall promptly return to the other Party all documents, materials, and other property containing or reflecting the other Party's Confidential Information and shall certify in writing that it has complied with this requirement.",
        "The indemnifying Party shall defend, indemnify, and hold harmless the indemnified Party and its officers, directors, employees, agents, and successors from and against all losses, damages, liabilities, deficiencies, actions, judgments, interest, awards, penalties, fines, costs, or expenses of whatever kind.",
    ]
    for i in range(count):
        add_clause(doc, legal_paragraphs[i % len(legal_paragraphs)])


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Add footer with page numbers
    add_page_number_footer(section)

    # =================== TITLE PAGE (Page 1) ===================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = title.add_run('MASTER SERVICES AGREEMENT')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = subtitle.add_run('Between\nMeridian Technology Solutions, Inc.\nand\nPacific Northwest Healthcare Group')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_before = Pt(36)
    r = date_para.add_run('Effective Date: January 15, 2025')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    doc.add_page_break()

    # =================== TABLE OF CONTENTS (Page 2) ===================
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = toc_title.add_run('TABLE OF CONTENTS')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    toc_entries = [
        'Article I - Definitions',
        'Article II - Scope of Services',
        'Article III - Term and Termination',
        'Article IV - Compensation and Payment',
        'Article V - Representations and Warranties',
        'Article VI - Intellectual Property',
        'Article VII - Confidentiality',
        'Article VIII - Limitation of Liability',
        'Article IX - Indemnification',
        'Article X - Insurance Requirements',
        'Article XI - Compliance and Regulatory',
        'Article XII - Dispute Resolution',
        'Article XIII - General Provisions',
        'Exhibit A - Statement of Work',
        'Exhibit B - Service Level Agreement',
        'Exhibit C - Fee Schedule',
        'Exhibit D - Data Processing Addendum',
        'Exhibit E - Insurance Certificate Requirements',
    ]
    for entry in toc_entries:
        p = doc.add_paragraph()
        r = p.add_run(entry)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # =================== ARTICLES (Pages 3-15) ===================
    articles = [
        ('Article I - Definitions', [
            '1.1 Definitions. "Agreement" means this Master Services Agreement, including all Exhibits, Schedules, and Amendments hereto.',
            '1.2 Affiliate. "Affiliate" means any entity that directly or indirectly controls, is controlled by, or is under common control with a Party, where "control" means the ownership of more than fifty percent (50%) of the voting securities of such entity.',
            '1.3 Confidential Information. "Confidential Information" means all non-public information disclosed by one Party to the other, whether orally or in writing, that is designated as confidential or that reasonably should be understood to be confidential given the nature of the information and the circumstances of disclosure.',
            '1.4 Deliverables. "Deliverables" means any work product, materials, reports, analyses, software, documentation, or other tangible items to be provided by Service Provider to Client under a Statement of Work.',
            '1.5 Effective Date. "Effective Date" means January 15, 2025, the date on which this Agreement becomes binding upon the Parties.',
            '1.6 Intellectual Property Rights. "Intellectual Property Rights" means all rights in patents, copyrights, trademarks, trade secrets, and other proprietary rights, whether registered or unregistered, and all applications and registrations relating thereto.',
        ]),
        ('Article II - Scope of Services', [
            '2.1 Services. Service Provider shall provide to Client the services described in each Statement of Work executed by the Parties (the "Services"). Each Statement of Work shall describe in reasonable detail the specific Services to be performed, the Deliverables, the timeline for performance, and the fees payable.',
            '2.2 Standard of Care. Service Provider shall perform the Services in a professional and workmanlike manner, consistent with generally accepted industry standards and practices. Service Provider shall assign qualified personnel with the necessary skills, training, and experience to perform the Services.',
            '2.3 Subcontractors. Service Provider may engage subcontractors to perform portions of the Services, provided that Service Provider shall remain responsible for the performance of the Services by its subcontractors as if Service Provider had performed such Services itself. Service Provider shall ensure that each subcontractor agrees to be bound by confidentiality obligations no less restrictive than those in this Agreement.',
            '2.4 Client Cooperation. Client shall provide reasonable cooperation and assistance to Service Provider as reasonably necessary for Service Provider to perform the Services, including providing timely access to Client personnel, facilities, systems, data, and information.',
        ]),
        ('Article III - Term and Termination', [
            '3.1 Term. This Agreement shall commence on the Effective Date and shall continue for an initial term of three (3) years (the "Initial Term"), unless earlier terminated in accordance with this Article III. Upon expiration of the Initial Term, this Agreement shall automatically renew for successive one (1) year renewal terms.',
            '3.2 Termination for Convenience. Either Party may terminate this Agreement for any reason upon ninety (90) days prior written notice to the other Party.',
            '3.3 Termination for Cause. Either Party may terminate this Agreement immediately upon written notice if the other Party: (a) materially breaches this Agreement and fails to cure such breach within thirty (30) days after receiving written notice thereof; or (b) becomes insolvent, files for bankruptcy, or has a receiver appointed for a substantial part of its assets.',
            '3.4 Effect of Termination. Upon termination or expiration of this Agreement: (a) Service Provider shall cease performing the Services; (b) Client shall pay Service Provider for all Services performed and expenses incurred through the effective date of termination; and (c) each Party shall return or destroy all Confidential Information of the other Party.',
        ]),
        ('Article IV - Compensation and Payment', [
            '4.1 Fees. Client shall pay Service Provider the fees set forth in each Statement of Work and as further detailed in Exhibit C - Fee Schedule. Unless otherwise specified in a Statement of Work, Service Provider shall invoice Client monthly in arrears for Services performed during the preceding month.',
            '4.2 Payment Terms. Client shall pay all undisputed invoices within thirty (30) days of receipt. Late payments shall bear interest at the rate of one and one-half percent (1.5%) per month or the maximum rate permitted by law, whichever is less. Client shall reimburse Service Provider for all reasonable and documented out-of-pocket expenses incurred in connection with the performance of the Services.',
            '4.3 Taxes. The fees set forth herein are exclusive of all applicable taxes. Client shall be responsible for payment of all sales, use, value-added, excise, and other taxes imposed on the Services, excluding taxes based on Service Provider\'s net income.',
            '4.4 Disputed Invoices. If Client disputes any portion of an invoice in good faith, Client shall notify Service Provider in writing within fifteen (15) days of receipt of such invoice, specifying in reasonable detail the basis for the dispute. The Parties shall work together in good faith to resolve any such dispute.',
        ]),
        ('Article V - Representations and Warranties', [
            '5.1 Mutual Representations. Each Party represents and warrants that: (a) it is duly organized, validly existing, and in good standing under the laws of its jurisdiction of organization; (b) it has full power and authority to enter into this Agreement; (c) the execution and performance of this Agreement does not conflict with any other agreement to which it is a party.',
            '5.2 Service Provider Warranties. Service Provider represents and warrants that: (a) the Services shall be performed in a professional manner consistent with industry standards; (b) the Deliverables shall conform to the specifications set forth in the applicable Statement of Work; (c) to Service Provider\'s knowledge, the Deliverables shall not infringe any third-party Intellectual Property Rights.',
            '5.3 Disclaimer. EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, NEITHER PARTY MAKES ANY WARRANTIES OF ANY KIND, WHETHER EXPRESS, IMPLIED, STATUTORY, OR OTHERWISE, INCLUDING ANY WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, OR NON-INFRINGEMENT.',
        ]),
        ('Article VI - Intellectual Property', [
            '6.1 Pre-Existing IP. Each Party shall retain all rights, title, and interest in its pre-existing Intellectual Property Rights. Neither Party grants the other any rights in its pre-existing IP except as expressly provided herein.',
            '6.2 Work Product. Unless otherwise specified in a Statement of Work, all Deliverables and work product created by Service Provider specifically for Client under this Agreement shall be considered "works made for hire" and shall be owned by Client. To the extent any Deliverable does not qualify as a work made for hire, Service Provider hereby assigns to Client all right, title, and interest in such Deliverable.',
            '6.3 License to Service Provider Tools. Service Provider grants Client a non-exclusive, perpetual, royalty-free license to use any Service Provider tools, methodologies, or frameworks incorporated into the Deliverables, solely to the extent necessary for Client to use the Deliverables for their intended purpose.',
        ]),
        ('Article VII - Confidentiality', [
            '7.1 Obligations. The Receiving Party shall: (a) hold all Confidential Information in strict confidence; (b) not disclose Confidential Information to any third party except as permitted herein; (c) use Confidential Information only for the purposes of this Agreement; and (d) protect Confidential Information using the same degree of care it uses to protect its own confidential information, but in no event less than reasonable care.',
            '7.2 Permitted Disclosures. The Receiving Party may disclose Confidential Information: (a) to its employees, agents, and subcontractors who have a need to know and who are bound by confidentiality obligations no less restrictive than those herein; and (b) as required by law, regulation, or court order, provided the Receiving Party gives prompt notice to the Disclosing Party.',
            '7.3 Duration. The obligations under this Article VII shall survive the termination or expiration of this Agreement for a period of five (5) years.',
        ]),
        ('Article VIII - Limitation of Liability', [
            '8.1 Limitation. EXCEPT FOR BREACHES OF CONFIDENTIALITY, INDEMNIFICATION OBLIGATIONS, OR WILLFUL MISCONDUCT, IN NO EVENT SHALL EITHER PARTY\'S AGGREGATE LIABILITY UNDER THIS AGREEMENT EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT TO SERVICE PROVIDER DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY PRECEDING THE EVENT GIVING RISE TO SUCH LIABILITY.',
            '8.2 Exclusion of Consequential Damages. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, CONSEQUENTIAL, SPECIAL, OR EXEMPLARY DAMAGES, INCLUDING DAMAGES FOR LOSS OF PROFITS, GOODWILL, DATA, OR OTHER INTANGIBLE LOSSES, EVEN IF SUCH PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.',
        ]),
        ('Article IX - Indemnification', [
            '9.1 Service Provider Indemnification. Service Provider shall defend, indemnify, and hold harmless Client and its officers, directors, employees, and agents from and against any and all third-party claims, losses, damages, liabilities, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to: (a) any negligent or wrongful act or omission of Service Provider in the performance of the Services; (b) any breach of Service Provider\'s representations, warranties, or obligations under this Agreement; or (c) any claim that the Deliverables infringe any third-party Intellectual Property Rights.',
            '9.2 Client Indemnification. Client shall defend, indemnify, and hold harmless Service Provider from any third-party claims arising out of: (a) Client\'s use of the Deliverables in a manner not authorized by this Agreement; (b) any materials or data provided by Client that infringe third-party rights; or (c) any breach of Client\'s obligations under this Agreement.',
            '9.3 Indemnification Procedure. The indemnified Party shall: (a) promptly notify the indemnifying Party in writing of any claim; (b) grant the indemnifying Party sole control of the defense and settlement; and (c) provide reasonable cooperation at the indemnifying Party\'s expense.',
        ]),
        ('Article X - Insurance Requirements', [
            '10.1 Required Coverage. Service Provider shall maintain throughout the term of this Agreement: (a) commercial general liability insurance with limits of not less than $2,000,000 per occurrence and $5,000,000 in the aggregate; (b) professional liability (errors and omissions) insurance with limits of not less than $3,000,000 per claim; (c) workers\' compensation insurance as required by applicable law; and (d) cyber liability insurance with limits of not less than $5,000,000 per occurrence.',
            '10.2 Evidence of Insurance. Service Provider shall provide Client with certificates of insurance evidencing the required coverage upon request and shall provide thirty (30) days\' advance written notice of any material change or cancellation of such insurance. The required insurance certificate specifications are detailed in Exhibit E.',
        ]),
        ('Article XI - Compliance and Regulatory', [
            '11.1 Compliance with Laws. Each Party shall comply with all applicable federal, state, and local laws, regulations, and ordinances in the performance of its obligations under this Agreement, including without limitation the Health Insurance Portability and Accountability Act (HIPAA), the Health Information Technology for Economic and Clinical Health Act (HITECH), and any implementing regulations.',
            '11.2 Data Protection. Service Provider shall implement and maintain appropriate technical and organizational measures to protect Client data in accordance with the Data Processing Addendum attached as Exhibit D. Service Provider shall promptly notify Client of any actual or suspected data breach affecting Client data.',
            '11.3 Audit Rights. Client shall have the right to audit Service Provider\'s compliance with this Agreement upon thirty (30) days\' prior written notice, no more than once per calendar year. Service Provider shall cooperate fully with any such audit.',
        ]),
        ('Article XII - Dispute Resolution', [
            '12.1 Negotiation. Any dispute arising out of or relating to this Agreement shall first be submitted to the senior management of each Party for resolution through good faith negotiation. The Parties shall use commercially reasonable efforts to resolve any dispute within thirty (30) days.',
            '12.2 Mediation. If the dispute is not resolved through negotiation, either Party may initiate mediation by delivering written notice to the other Party. The mediation shall be conducted under the rules of the American Arbitration Association in Wilmington, Delaware.',
            '12.3 Arbitration. If mediation fails to resolve the dispute within sixty (60) days, the dispute shall be resolved by binding arbitration conducted under the Commercial Arbitration Rules of the American Arbitration Association. The arbitration shall be conducted by a panel of three (3) arbitrators in Wilmington, Delaware.',
        ]),
        ('Article XIII - General Provisions', [
            '13.1 Entire Agreement. This Agreement, including all Exhibits, Schedules, and Statements of Work, constitutes the entire agreement between the Parties with respect to the subject matter hereof and supersedes all prior and contemporaneous agreements, understandings, negotiations, and discussions.',
            '13.2 Amendment. This Agreement may not be amended or modified except by a written instrument signed by both Parties.',
            '13.3 Governing Law. This Agreement shall be governed by the laws of the State of Delaware, without regard to its conflict of laws provisions.',
            '13.4 Force Majeure. Neither Party shall be liable for any failure or delay in performing its obligations under this Agreement to the extent that such failure or delay results from circumstances beyond the reasonable control of such Party, including acts of God, natural disasters, war, terrorism, labor disputes, or government actions.',
            '13.5 Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original, and all of which together shall constitute one and the same instrument.',
        ]),
    ]

    for article_title, clauses in articles:
        add_heading(doc, article_title, level=1)
        for clause_text in clauses:
            add_clause(doc, clause_text, bold_first_sentence=True)
        # Add filler text to ensure we fill enough pages
        add_filler_text(doc, count=2)

    # =================== SIGNATURE PAGE (Page ~15) ===================
    doc.add_page_break()
    sig_title = doc.add_paragraph()
    sig_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = sig_title.add_run('SIGNATURE PAGE')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    add_clause(doc, 'IN WITNESS WHEREOF, the Parties have executed this Master Services Agreement as of the Effective Date first written above.')

    doc.add_paragraph()
    for company, name, title in [
        ('MERIDIAN TECHNOLOGY SOLUTIONS, INC.', 'Alexandra M. Richardson', 'Chief Executive Officer'),
        ('PACIFIC NORTHWEST HEALTHCARE GROUP', 'Dr. Robert J. Kawamura', 'President and Chief Operating Officer'),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(company)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run('By: ___________________________________')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        r = p.add_run(f'Name: {name}')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        r = p.add_run(f'Title: {title}')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        r = p.add_run('Date: January 15, 2025')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        doc.add_paragraph()

    # =================== EXHIBITS (Pages 16-22) ===================
    # In the initial state, exhibits continue with the same section
    # (continuous page numbering, no restart, no A- prefix)
    doc.add_page_break()

    add_heading(doc, 'EXHIBITS', level=0)

    # Exhibit A - Statement of Work
    add_heading(doc, 'Exhibit A - Statement of Work', level=1)
    add_clause(doc, '1. Project Overview. Meridian Technology Solutions, Inc. ("Service Provider") shall provide Pacific Northwest Healthcare Group ("Client") with a comprehensive Electronic Health Records (EHR) system modernization project, including system design, development, testing, deployment, and post-deployment support services.')
    add_clause(doc, '2. Project Scope. The Services shall include: (a) analysis of Client\'s current EHR systems and workflows; (b) design and development of a new cloud-based EHR platform; (c) data migration from legacy systems; (d) integration with existing hospital information systems; (e) end-user training and documentation; and (f) twelve (12) months of post-deployment technical support.')
    add_clause(doc, '3. Timeline. The project shall be completed in four phases over an eighteen (18) month period: Phase 1 - Discovery and Planning (3 months); Phase 2 - Design and Development (6 months); Phase 3 - Testing and Deployment (6 months); Phase 4 - Stabilization and Support (3 months).')
    add_clause(doc, '4. Key Personnel. Service Provider shall assign the following key personnel to the project: Project Director - Jennifer Walsh; Technical Lead - David Chen; Integration Architect - Priya Ramanathan; Training Coordinator - Michael Torres.')
    add_filler_text(doc, count=4)

    doc.add_page_break()

    # Exhibit B - Service Level Agreement
    add_heading(doc, 'Exhibit B - Service Level Agreement', level=1)
    add_clause(doc, '1. Service Availability. Service Provider shall ensure that the EHR platform maintains a minimum uptime of 99.95% measured on a monthly basis, excluding scheduled maintenance windows. Scheduled maintenance shall occur only during pre-approved maintenance windows (Sundays 2:00 AM - 6:00 AM Pacific Time).')
    add_clause(doc, '2. Response Times. Service Provider shall respond to and resolve support incidents within the following timeframes: Critical (P1) - System Down: Response within 15 minutes, Resolution within 4 hours; High (P2) - Major Feature Impaired: Response within 1 hour, Resolution within 8 hours; Medium (P3) - Minor Feature Impaired: Response within 4 hours, Resolution within 24 hours; Low (P4) - Information Request: Response within 8 hours, Resolution within 72 hours.')
    add_clause(doc, '3. Performance Metrics. The EHR platform shall meet the following performance benchmarks: (a) page load time not to exceed 2 seconds for 95% of transactions; (b) report generation not to exceed 30 seconds for standard reports; (c) search results returned within 3 seconds; (d) concurrent user support for a minimum of 2,500 simultaneous users.')
    add_clause(doc, '4. Service Credits. In the event Service Provider fails to meet the service levels set forth herein, Client shall be entitled to service credits calculated as follows: 99.9% - 99.95% uptime: 5% credit on monthly fees; 99.0% - 99.9% uptime: 10% credit; Below 99.0% uptime: 20% credit plus right to terminate for cause.')
    add_filler_text(doc, count=3)

    doc.add_page_break()

    # Exhibit C - Fee Schedule
    add_heading(doc, 'Exhibit C - Fee Schedule', level=1)
    add_clause(doc, '1. Professional Services Fees. The following hourly rates shall apply: Project Director - $350/hour; Technical Lead - $275/hour; Senior Developer - $225/hour; Developer - $175/hour; QA Engineer - $165/hour; Business Analyst - $195/hour; Training Specialist - $150/hour.')
    add_clause(doc, '2. Fixed Fee Components. Phase 1 - Discovery and Planning: $485,000; Phase 2 - Design and Development: $2,150,000; Phase 3 - Testing and Deployment: $1,275,000; Phase 4 - Stabilization and Support: $390,000. Total Fixed Fee: $4,300,000.')
    add_clause(doc, '3. Recurring Fees. Monthly hosting and infrastructure: $45,000/month; Annual maintenance and support (Year 2+): $515,000/year; Additional storage (per TB): $2,500/month.')
    add_clause(doc, '4. Expense Policy. Reasonable travel expenses shall be reimbursed at cost with prior approval. Air travel shall be at coach/economy class rates. Hotel accommodations shall not exceed $250 per night. Meals shall not exceed $75 per day.')
    add_filler_text(doc, count=3)

    doc.add_page_break()

    # Exhibit D - Data Processing Addendum
    add_heading(doc, 'Exhibit D - Data Processing Addendum', level=1)
    add_clause(doc, '1. Data Processing. Service Provider shall process Client data solely for the purpose of providing the Services under this Agreement. Service Provider shall not sell, share, or use Client data for any purpose other than performing its obligations hereunder.')
    add_clause(doc, '2. Security Measures. Service Provider shall implement and maintain: (a) encryption of data at rest (AES-256) and in transit (TLS 1.3); (b) multi-factor authentication for all administrative access; (c) intrusion detection and prevention systems; (d) regular vulnerability assessments and penetration testing; (e) security incident and event management (SIEM) monitoring.')
    add_clause(doc, '3. Data Breach Notification. Service Provider shall notify Client within twenty-four (24) hours of discovering any actual or reasonably suspected data breach affecting Client data. Such notification shall include the nature of the breach, the categories and approximate number of records affected, and the measures taken to mitigate the breach.')
    add_clause(doc, '4. Data Retention and Deletion. Upon termination of this Agreement, Service Provider shall, at Client\'s election, return all Client data in a standard, machine-readable format or securely delete all Client data within thirty (30) days and provide written certification of such deletion.')
    add_filler_text(doc, count=3)

    doc.add_page_break()

    # Exhibit E - Insurance Certificate Requirements
    add_heading(doc, 'Exhibit E - Insurance Certificate Requirements', level=1)
    add_clause(doc, '1. Certificate Specifications. All certificates of insurance shall: (a) name Pacific Northwest Healthcare Group as an additional insured; (b) include a waiver of subrogation in favor of Client; (c) state that coverage shall not be cancelled or materially changed without thirty (30) days\' prior written notice to Client; (d) be issued by insurers rated A- VII or better by A.M. Best.')
    add_clause(doc, '2. Coverage Details. The certificate shall evidence the following minimum coverage: Commercial General Liability ($2M/$5M); Professional Liability ($3M per claim); Workers\' Compensation (statutory limits); Cyber Liability ($5M per occurrence); Commercial Auto Liability ($1M combined single limit); Umbrella/Excess Liability ($10M).')
    add_clause(doc, '3. Annual Renewal. Service Provider shall provide updated certificates of insurance to Client within thirty (30) days of each policy renewal date. Failure to maintain the required insurance coverage shall constitute a material breach of this Agreement.')
    add_filler_text(doc, count=4)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
