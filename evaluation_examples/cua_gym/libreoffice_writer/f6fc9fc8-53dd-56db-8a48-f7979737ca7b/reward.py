"""
Reward Script: Page numbering with Arabic numerals for main body and 'A-[number]' for exhibits
Task ID: writer_legal_044
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Document has >= 2 sections (section break exists)
  Component 2 (0.20): Section break is positioned near the Exhibits section
  Component 3 (0.25): Second section restarts page numbering at 1
  Component 4 (0.30): Second section footer contains 'A-' prefix before PAGE field
"""

import os
from docx import Document
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_044'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice Writer."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)

    # Component 1: Document has >= 2 sections (0.25 points)
    # The task requires a section break between the main body and exhibits.
    # Initial doc has only 1 section, so this differentiates initial from golden.
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 -- Document has {num_sections} sections (>= 2) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 -- Document has {num_sections} section(s), expected >= 2")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Section break is positioned near the Exhibits section (0.20 points)
    # The section break should be at or just before the paragraph containing "EXHIBITS".
    # We check that a paragraph with "EXHIBITS" text exists in the second section's range.
    try:
        if num_sections >= 2:
            # Find the paragraph index where the section break (sectPr in pPr) is
            break_para_idx = None
            for i, para in enumerate(doc.paragraphs):
                pPr = para._element.find('w:pPr', NS)
                if pPr is not None:
                    sectPr_in_pPr = pPr.find('w:sectPr', NS)
                    if sectPr_in_pPr is not None:
                        break_para_idx = i
                        break

            if break_para_idx is not None:
                # Check that within a few paragraphs after the break, we find "EXHIBITS"
                exhibits_nearby = False
                for j in range(max(0, break_para_idx - 2), min(len(doc.paragraphs), break_para_idx + 5)):
                    if 'EXHIBIT' in doc.paragraphs[j].text.upper():
                        exhibits_nearby = True
                        break

                if exhibits_nearby:
                    print(f"PASS: Component 2 -- Section break at para {break_para_idx}, near Exhibits section (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"FAIL: Component 2 -- Section break at para {break_para_idx}, but no 'EXHIBITS' text nearby")
            else:
                print(f"FAIL: Component 2 -- No section break found in paragraph pPr elements")
        else:
            print(f"FAIL: Component 2 -- Only 1 section, no section break to check")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Second section restarts page numbering at 1 (0.25 points)
    # The pgNumType element with start="1" must be in the second section's sectPr.
    try:
        if num_sections >= 2:
            sec1 = doc.sections[1]
            sectPr = sec1._sectPr
            pgNumType_list = sectPr.findall('w:pgNumType', NS)
            restart_found = False
            for pnt in pgNumType_list:
                start_val = pnt.get(f'{{{NS["w"]}}}start')
                if start_val == '1':
                    restart_found = True
                    break

            if restart_found:
                print(f"PASS: Component 3 -- Section 2 has pgNumType start='1' (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 -- Section 2 pgNumType missing or start != '1'. Found: {[dict(pnt.attrib) for pnt in pgNumType_list]}")
        else:
            print(f"FAIL: Component 3 -- Only 1 section, cannot check page numbering restart")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Second section footer has 'A-' prefix before PAGE field (0.30 points)
    # The footer of section 2 should contain literal text "A-" followed by a PAGE field code.
    try:
        if num_sections >= 2:
            sec1 = doc.sections[1]
            footer = sec1.footer

            # Check that footer is not linked to previous (has its own content)
            if footer.is_linked_to_previous:
                print(f"FAIL: Component 4 -- Section 2 footer is linked to previous (should have its own)")
            else:
                # Look for "A-" text in footer runs, and a PAGE field code
                has_a_prefix = False
                has_page_field = False

                for para in footer.paragraphs:
                    # Check all runs for "A-" text
                    for run in para.runs:
                        if 'A-' in run.text:
                            has_a_prefix = True
                        # Check for PAGE instrText field code
                        instr_texts = run._element.findall('.//w:instrText', NS)
                        for instr in instr_texts:
                            if instr.text and 'PAGE' in instr.text:
                                has_page_field = True

                    # Also check via XML for PAGE field in case runs don't capture it
                    para_xml = etree.tostring(para._element, encoding='unicode')
                    if 'PAGE' in para_xml and 'instrText' in para_xml:
                        has_page_field = True

                if has_a_prefix and has_page_field:
                    print(f"PASS: Component 4 -- Section 2 footer has 'A-' prefix and PAGE field (0.30 pts)")
                    total_score += 0.30
                elif has_a_prefix:
                    print(f"PARTIAL: Component 4 -- Has 'A-' prefix but no PAGE field (0.15 pts)")
                    total_score += 0.15
                elif has_page_field:
                    print(f"PARTIAL: Component 4 -- Has PAGE field but no 'A-' prefix (0.10 pts)")
                    total_score += 0.10
                else:
                    print(f"FAIL: Component 4 -- Section 2 footer missing both 'A-' prefix and PAGE field")
        else:
            print(f"FAIL: Component 4 -- Only 1 section, cannot check footer")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
