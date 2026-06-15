"""
Reward Script: Professional letterhead template verification
Task ID: writer_pd_042
Domain: libreoffice_writer
Scoring:
  Component 1 — Header has active table structure (0.15)
  Component 2 — Logo placeholder image ~3cm x 1.5cm, blue (0.20)
  Component 3 — Company name centered, bold, 16pt (0.25)
  Component 4 — Contact info right-aligned, 8pt (0.20)
  Component 5 — Footer with top border line, centered address in 8pt gray (0.20)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.oxml.ns import qn
from io import BytesIO

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_042'


def verify_task(file_path):
    """
    Verify letterhead template completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    header = section.header
    footer = section.footer

    # =========================================================================
    # Component 1: Header has an active table structure (0.15 points)
    # The golden header uses a 3-column table for layout.
    # Initial state has header linked_to_previous=True with no table.
    # =========================================================================
    try:
        header_elements = [child.tag.split('}')[-1] for child in header._element]
        has_table = 'tbl' in header_elements
        is_linked = header.is_linked_to_previous

        if not is_linked and has_table:
            # Count table columns
            tbl = header._element.findall(qn('w:tbl'))
            if tbl:
                grid_cols = tbl[0].findall(qn('w:tblGrid') + '/' + qn('w:gridCol'))
                col_count = len(grid_cols)
                if col_count >= 3:
                    print(f"PASS: Component 1 — Header has active table with {col_count} columns (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"PARTIAL: Component 1 — Header table has only {col_count} columns, expected >= 3")
                    total_score += 0.05
            else:
                print("FAIL: Component 1 — Header element has 'tbl' tag but no w:tbl found")
        elif not is_linked:
            # Header is active but uses tab stops or other layout instead of table
            # Give partial credit if header is at least not linked
            header_text = ' '.join(p.text for p in header.paragraphs).strip()
            if header_text:
                print(f"PARTIAL: Component 1 — Header is active with text but no table layout (0.05 pts)")
                total_score += 0.05
            else:
                print("FAIL: Component 1 — Header is active but empty, no table structure")
        else:
            print("FAIL: Component 1 — Header is still linked to previous (empty/default)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Logo placeholder image ~3cm x 1.5cm, blue (0.20 points)
    # The golden header left cell has a blue rectangle image (1080000 x 540000 EMU).
    # 3cm = 1080000 EMU, 1.5cm = 540000 EMU
    # Initial state has no images in header.
    # =========================================================================
    try:
        # Check for inline drawing elements in header XML
        header_xml = header._element.xml
        has_drawing = '<w:drawing>' in header_xml or 'w:drawing' in header_xml

        # Also check for images via relationships
        header_images = []
        for rel in header.part.rels.values():
            if 'image' in rel.reltype:
                header_images.append(rel)

        if has_drawing and len(header_images) > 0:
            # Check image dimensions from XML
            ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            extents = header._element.findall(f'.//{{{ns_wp}}}extent')

            image_ok = False
            for ext in extents:
                cx = int(ext.get('cx', 0))
                cy = int(ext.get('cy', 0))
                # 3cm = 1080000 EMU, 1.5cm = 540000 EMU — allow 20% tolerance
                cx_ok = abs(cx - 1080000) < 1080000 * 0.25
                cy_ok = abs(cy - 540000) < 540000 * 0.25
                if cx_ok and cy_ok:
                    image_ok = True
                    print(f"  Image dimensions: cx={cx} ({cx/360000:.1f}cm), cy={cy} ({cy/360000:.1f}cm)")

            if image_ok:
                # Check if image is blue-ish
                try:
                    from PIL import Image
                    import numpy as np
                    img_blob = header_images[0].target_part.blob
                    img = Image.open(BytesIO(img_blob))
                    arr = np.array(img)
                    mean_color = arr.mean(axis=(0, 1))
                    # Blue channel should dominate
                    is_blue = mean_color[2] > 100 and mean_color[2] > mean_color[0] * 2
                    if is_blue:
                        print(f"PASS: Component 2 — Blue logo placeholder with correct dimensions (0.20 pts)")
                        total_score += 0.20
                    else:
                        print(f"PARTIAL: Component 2 — Logo image correct size but not blue (mean RGB={mean_color[:3]}) (0.10 pts)")
                        total_score += 0.10
                except Exception as img_e:
                    # Can't verify color, give credit for dimensions
                    print(f"PARTIAL: Component 2 — Logo image correct size, color check failed: {img_e} (0.12 pts)")
                    total_score += 0.12
            else:
                # Image exists but wrong dimensions
                dims = [(int(e.get('cx', 0)), int(e.get('cy', 0))) for e in extents]
                print(f"PARTIAL: Component 2 — Logo image found but wrong dimensions: {dims} (0.08 pts)")
                total_score += 0.08
        else:
            print(f"FAIL: Component 2 — No logo image in header (drawing={has_drawing}, images={len(header_images)})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: Company name "Apex Global Partners" centered, bold, 16pt (0.25 points)
    # In golden: center cell of header table has "Apex Global Partners" with w:b and w:sz=32 (16pt)
    # Initial state: no such text anywhere in header.
    # =========================================================================
    try:
        found_name = False
        name_bold = False
        name_size_ok = False
        name_centered = False

        # Search all paragraphs in header (including table cells)
        all_header_paras = []
        # Direct header paragraphs
        all_header_paras.extend(header.paragraphs)
        # Table cell paragraphs
        for tbl_el in header._element.findall(qn('w:tbl')):
            for tc in tbl_el.findall('.//' + qn('w:tc')):
                for p_el in tc.findall(qn('w:p')):
                    from docx.text.paragraph import Paragraph
                    all_header_paras.append(Paragraph(p_el, header))

        for para in all_header_paras:
            full_text = para.text.strip()
            if 'Apex Global Partners' in full_text:
                found_name = True
                # Check alignment
                jc = para._element.find(qn('w:pPr') + '/' + qn('w:jc'))
                if jc is not None and jc.get(qn('w:val')) == 'center':
                    name_centered = True
                elif para.paragraph_format.alignment is not None:
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        name_centered = True

                for run in para.runs:
                    if 'Apex Global Partners' in run.text:
                        # Check bold
                        b_el = run._element.find(qn('w:rPr') + '/' + qn('w:b'))
                        if b_el is not None or run.font.bold:
                            name_bold = True
                        # Check size (16pt = sz val 32 in half-points)
                        sz_el = run._element.find(qn('w:rPr') + '/' + qn('w:sz'))
                        if sz_el is not None:
                            sz_val = int(sz_el.get(qn('w:val'), '0'))
                            if 28 <= sz_val <= 36:  # ~14pt to ~18pt tolerance
                                name_size_ok = True
                        elif run.font.size is not None:
                            pt_val = run.font.size.pt
                            if 14 <= pt_val <= 18:
                                name_size_ok = True

        if found_name:
            sub_score = 0.0
            if name_centered:
                sub_score += 0.08
            if name_bold:
                sub_score += 0.08
            if name_size_ok:
                sub_score += 0.09

            if sub_score >= 0.25:
                print(f"PASS: Component 3 — 'Apex Global Partners' centered={name_centered}, bold={name_bold}, 16pt={name_size_ok} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"PARTIAL: Component 3 — 'Apex Global Partners' found but: centered={name_centered}, bold={name_bold}, 16pt={name_size_ok} ({sub_score} pts)")
                total_score += sub_score
        else:
            print("FAIL: Component 3 — 'Apex Global Partners' not found in header")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================================
    # Component 4: Contact info right-aligned, 8pt (0.20 points)
    # Golden: right cell has phone, email, website in 8pt (sz=16), right-aligned
    # Initial: no contact info in header.
    # =========================================================================
    try:
        found_phone = False
        found_email = False
        found_website = False
        contact_right = False
        contact_8pt = False

        # Search header paragraphs (including table cells)
        for para in all_header_paras:
            full_text = para.text.strip().lower()
            has_contact = False
            if 'phone' in full_text or '+1' in full_text or '555' in full_text:
                found_phone = True
                has_contact = True
            if '@' in full_text or 'email' in full_text:
                found_email = True
                has_contact = True
            if 'www.' in full_text or '.com' in full_text:
                found_website = True
                has_contact = True

            if has_contact:
                # Check right alignment
                jc = para._element.find(qn('w:pPr') + '/' + qn('w:jc'))
                if jc is not None and jc.get(qn('w:val')) == 'right':
                    contact_right = True

                # Check font size (8pt = sz val 16 in half-points)
                for run in para.runs:
                    sz_el = run._element.find(qn('w:rPr') + '/' + qn('w:sz'))
                    if sz_el is not None:
                        sz_val = int(sz_el.get(qn('w:val'), '0'))
                        if 14 <= sz_val <= 18:  # ~7pt to ~9pt tolerance
                            contact_8pt = True
                    elif run.font.size is not None:
                        if 7 <= run.font.size.pt <= 9:
                            contact_8pt = True

        contact_count = sum([found_phone, found_email, found_website])
        if contact_count >= 2:
            sub_score = 0.0
            # Points for having contact items
            sub_score += min(contact_count, 3) * 0.04  # up to 0.12
            if contact_right:
                sub_score += 0.04
            if contact_8pt:
                sub_score += 0.04

            if contact_count >= 3 and contact_right and contact_8pt:
                print(f"PASS: Component 4 — Contact info (phone={found_phone}, email={found_email}, web={found_website}), right-aligned, 8pt (0.20 pts)")
                total_score += 0.20
            else:
                actual_pts = min(sub_score, 0.20)
                print(f"PARTIAL: Component 4 — phone={found_phone}, email={found_email}, web={found_website}, right={contact_right}, 8pt={contact_8pt} ({actual_pts} pts)")
                total_score += actual_pts
        elif contact_count == 1:
            print(f"PARTIAL: Component 4 — Only {contact_count}/3 contact items found (0.04 pts)")
            total_score += 0.04
        else:
            print("FAIL: Component 4 — No contact information found in header")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =========================================================================
    # Component 5: Footer with top border line, centered address in 8pt gray (0.20 points)
    # Golden: footer para has w:pBdr/w:top (thin line), centered, text in 8pt gray (#808080)
    # Initial: footer is empty/linked.
    # =========================================================================
    try:
        footer_text = ''
        footer_centered = False
        footer_has_border = False
        footer_gray = False
        footer_8pt = False

        for para in footer.paragraphs:
            text = para.text.strip()
            if text:
                footer_text = text

                # Check centered alignment
                jc = para._element.find(qn('w:pPr') + '/' + qn('w:jc'))
                if jc is not None and jc.get(qn('w:val')) == 'center':
                    footer_centered = True

                # Check top border
                top_bdr = para._element.find(qn('w:pPr') + '/' + qn('w:pBdr') + '/' + qn('w:top'))
                if top_bdr is not None:
                    bdr_val = top_bdr.get(qn('w:val'), '')
                    if bdr_val != 'none':
                        footer_has_border = True

                # Check runs for gray color and 8pt size
                for run in para.runs:
                    # Color check
                    color_el = run._element.find(qn('w:rPr') + '/' + qn('w:color'))
                    if color_el is not None:
                        color_val = color_el.get(qn('w:val'), '').upper()
                        # Accept gray-ish colors (808080, 7F7F7F, 999999, etc.)
                        if color_val in ('808080', '7F7F7F', '999999', 'A0A0A0', 'GRAY', 'GREY'):
                            footer_gray = True
                        elif run.font.color.rgb is not None:
                            r, g, b = run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2]
                            if 80 <= r <= 180 and abs(r - g) < 30 and abs(g - b) < 30:
                                footer_gray = True
                    elif run.font.color.rgb is not None:
                        r, g, b = run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2]
                        if 80 <= r <= 180 and abs(r - g) < 30 and abs(g - b) < 30:
                            footer_gray = True

                    # Size check
                    sz_el = run._element.find(qn('w:rPr') + '/' + qn('w:sz'))
                    if sz_el is not None:
                        sz_val = int(sz_el.get(qn('w:val'), '0'))
                        if 14 <= sz_val <= 18:
                            footer_8pt = True
                    elif run.font.size is not None:
                        if 7 <= run.font.size.pt <= 9:
                            footer_8pt = True

        has_address = bool(footer_text) and ('avenue' in footer_text.lower() or 'street' in footer_text.lower() or 'suite' in footer_text.lower() or 'ny' in footer_text.lower())

        if has_address:
            sub_score = 0.06  # Base for having address text
            if footer_centered:
                sub_score += 0.04
            if footer_has_border:
                sub_score += 0.04
            if footer_gray:
                sub_score += 0.03
            if footer_8pt:
                sub_score += 0.03

            if footer_centered and footer_has_border and footer_gray and footer_8pt:
                print(f"PASS: Component 5 — Footer: '{footer_text[:50]}', centered, top border, 8pt gray (0.20 pts)")
                total_score += 0.20
            else:
                actual_pts = min(sub_score, 0.20)
                print(f"PARTIAL: Component 5 — Footer text found: centered={footer_centered}, border={footer_has_border}, gray={footer_gray}, 8pt={footer_8pt} ({actual_pts} pts)")
                total_score += actual_pts
        elif footer_text:
            print(f"PARTIAL: Component 5 — Footer has text '{footer_text[:40]}' but no address keywords (0.04 pts)")
            total_score += 0.04
        else:
            print("FAIL: Component 5 — Footer is empty, no address text")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/Letterhead_Template.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
