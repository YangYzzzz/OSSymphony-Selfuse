"""
Initial Setup: Insert a text box with 2 columns on page 1
Task ID: writer_obj_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'  # VM path — scripts run on the VM
TASK_ID = 'column_layout'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("Quarterly Business Review: Technology Division")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Calibri"

    doc.add_paragraph()

    # Introduction paragraph
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "This report summarizes the Technology Division's performance during the past quarter, "
        "highlighting key achievements, ongoing initiatives, and strategic priorities for the "
        "upcoming period. The division has made significant strides in infrastructure modernization, "
        "software development productivity, and cross-functional collaboration."
    )
    intro_run.font.size = Pt(11)
    intro_run.font.name = "Calibri"

    doc.add_paragraph()

    # Section heading
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("Executive Summary")
    h1_run.bold = True
    h1_run.font.size = Pt(13)
    h1_run.font.name = "Calibri"

    p1 = doc.add_paragraph()
    p1_run = p1.add_run(
        "The Technology Division closed Q1 2025 with a 14% increase in on-time project delivery "
        "compared to Q4 2024. Cloud migration efforts for core business applications reached 78% "
        "completion, ahead of the projected 70% target. Security incident response times improved "
        "by 32% following the rollout of the new automated monitoring platform."
    )
    p1_run.font.size = Pt(11)
    p1_run.font.name = "Calibri"

    doc.add_paragraph()

    # Key Highlights heading
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("Key Highlights")
    h2_run.bold = True
    h2_run.font.size = Pt(13)
    h2_run.font.name = "Calibri"

    highlights = [
        "Deployed 23 new microservices to production, reducing API latency by an average of 18%.",
        "Onboarded 5 new enterprise clients onto the SaaS platform, bringing total active accounts to 312.",
        "Completed Phase 2 of the data warehouse consolidation, enabling real-time analytics dashboards.",
        "Reduced average ticket resolution time from 4.2 days to 2.8 days through process automation.",
        "Launched the internal Developer Experience Portal, now used by 94% of engineering staff daily.",
    ]
    for item in highlights:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet_run = bullet.add_run(item)
        bullet_run.font.size = Pt(11)
        bullet_run.font.name = "Calibri"

    doc.add_paragraph()

    # Infrastructure heading
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("Infrastructure & Operations")
    h3_run.bold = True
    h3_run.font.size = Pt(13)
    h3_run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        "Infrastructure reliability remained above 99.9% SLA for the third consecutive quarter. "
        "The operations team successfully completed the transition to the new Tier-3 data center "
        "in Singapore, providing redundant failover capacity for Asia-Pacific customers. "
        "Network throughput capacity was upgraded from 40Gbps to 100Gbps across primary backbone links."
    )
    p2_run.font.size = Pt(11)
    p2_run.font.name = "Calibri"

    doc.add_paragraph()

    # Software Development heading
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("Software Development")
    h4_run.bold = True
    h4_run.font.size = Pt(13)
    h4_run.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3_run = p3.add_run(
        "Agile teams across the division completed a total of 847 story points this quarter, "
        "representing a 9% increase from Q4 2024. Code review turnaround times were reduced by "
        "introducing automated static analysis, catching 1,243 potential issues before human review. "
        "Unit test coverage across all active repositories improved from 61% to 74%."
    )
    p3_run.font.size = Pt(11)
    p3_run.font.name = "Calibri"

    # Save
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
