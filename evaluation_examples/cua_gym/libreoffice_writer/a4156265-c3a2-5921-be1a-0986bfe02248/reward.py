"""
Reward Script: Insert Next Page section breaks before each of 8 article headings
Task ID: writer_legal_069
Domain: libreoffice_writer
Scoring:
  Component 1: Correct number of sections (9 total) — 0.3 points
  Component 2: All 8 section breaks are nextPage type — 0.3 points
  Component 3: Each article heading starts at the beginning of its section — 0.4 points
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_069'

# The 8 article headings expected in the document
EXPECTED_ARTICLES = [
    'Article I',
    'Article II',
    'Article III',
    'Article IV',
    'Article V',
    'Article VI',
    'Article VII',
    'Article VIII',
]


def persist_app_state(domain):
    """Best-effort save via Ctrl+S in case doc is open in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that Next Page section breaks have been inserted before each
    of the 8 article headings in the lease agreement document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Identify article heading paragraph indices
    article_para_indices = []
    for i, para in enumerate(doc.paragraphs):
        for article_name in EXPECTED_ARTICLES:
            if para.text.strip().startswith(article_name):
                article_para_indices.append(i)
                break

    if len(article_para_indices) != 8:
        print(f"WARN: Expected 8 article headings, found {len(article_para_indices)}")

    # Count sections and collect section break info
    num_sections = len(doc.sections)

    # Collect all paragraphs that have sectPr (section breaks)
    # In docx, a section break in paragraph N means: everything up to and including
    # paragraph N is in that section, and paragraph N+1 starts a new section.
    paras_with_breaks = []
    for i, para in enumerate(doc.paragraphs):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            sect = pPr.find(qn('w:sectPr'))
            if sect is not None:
                type_el = sect.find(qn('w:type'))
                type_val = type_el.get(qn('w:val')) if type_el is not None else 'default'
                paras_with_breaks.append((i, type_val))

    print(f"INFO: Total sections: {num_sections}")
    print(f"INFO: Article heading paragraphs: {article_para_indices}")
    print(f"INFO: Paragraphs with section breaks: {paras_with_breaks}")

    # Component 1: Correct number of sections (9 total = 1 original + 8 breaks) — 0.3 points
    try:
        if num_sections >= 9:
            print(f"PASS: Component 1 — Document has {num_sections} sections (>= 9) (0.3 pts)")
            total_score += 0.3
        elif num_sections >= 5:
            partial = 0.3 * (num_sections - 1) / 8.0
            print(f"PARTIAL: Component 1 — Document has {num_sections} sections, expected 9 ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — Document has {num_sections} sections, expected 9")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All section breaks are nextPage type — 0.3 points
    try:
        next_page_count = sum(1 for _, t in paras_with_breaks if t == 'nextPage')
        total_breaks = len(paras_with_breaks)
        if next_page_count >= 8:
            print(f"PASS: Component 2 — {next_page_count} nextPage breaks found (>= 8) (0.3 pts)")
            total_score += 0.3
        elif next_page_count > 0:
            partial = 0.3 * next_page_count / 8.0
            print(f"PARTIAL: Component 2 — {next_page_count}/8 nextPage breaks ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No nextPage section breaks found (found {total_breaks} breaks of other types)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Section breaks are positioned before each article heading — 0.4 points
    # A section break in paragraph i means paragraph i+1 starts a new section.
    # So for each article heading at paragraph index h, we expect a break at paragraph h-1.
    try:
        break_para_set = set(idx for idx, t in paras_with_breaks if t == 'nextPage')
        correctly_placed = 0
        for h_idx in article_para_indices:
            expected_break_para = h_idx - 1
            if expected_break_para >= 0 and expected_break_para in break_para_set:
                correctly_placed += 1
                print(f"  OK: Article at para {h_idx} has nextPage break at para {expected_break_para}")
            else:
                print(f"  MISS: Article at para {h_idx} — no nextPage break at para {expected_break_para}")

        if correctly_placed == 8:
            print(f"PASS: Component 3 — All 8 article headings have correctly placed breaks (0.4 pts)")
            total_score += 0.4
        elif correctly_placed > 0:
            partial = 0.4 * correctly_placed / 8.0
            print(f"PARTIAL: Component 3 — {correctly_placed}/8 correctly placed breaks ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No article headings have correctly placed section breaks")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
