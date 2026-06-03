"""
Initial Setup: 8-article lease agreement flowing continuously without section breaks
Task ID: writer_legal_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Document title
    title = doc.add_heading('RESIDENTIAL LEASE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Preamble
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(12)
    run = preamble.add_run(
        'This Residential Lease Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between the Landlord and Tenant identified below. '
        'This Agreement sets forth the terms and conditions under which the Landlord '
        'agrees to lease the Premises to the Tenant, and the Tenant agrees to lease '
        'the Premises from the Landlord.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # --- Article I: Parties ---
    h1 = doc.add_heading('Article I: Parties', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'The Landlord, Harrington Property Management LLC, a limited liability company '
        'organized under the laws of the State of California, with its principal office '
        'located at 4200 Wilshire Boulevard, Suite 310, Los Angeles, California 90010, '
        'hereby agrees to lease the Premises described herein to the Tenant.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'The Tenant, Elena Vasquez-Morrison, an individual residing at 1847 Oakwood '
        'Drive, Apartment 12B, Pasadena, California 91104, hereby agrees to lease the '
        'Premises from the Landlord subject to the terms and conditions set forth in '
        'this Agreement. The Tenant acknowledges having read and understood all provisions '
        'contained herein prior to execution of this Agreement.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    # --- Article II: Premises ---
    doc.add_heading('Article II: Premises', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'The Landlord agrees to lease to the Tenant the residential property located at '
        '2738 Magnolia Court, Unit 5A, Santa Monica, California 90401 (the "Premises"). '
        'The Premises consist of a two-bedroom, one-bathroom apartment unit comprising '
        'approximately 1,150 square feet of living space, including a kitchen, living room, '
        'and a dedicated parking space designated as Space No. 47 in the building garage.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'The Premises are leased in their current condition as inspected by the Tenant on '
        'February 28, 2025. The Tenant acknowledges that the Premises are in satisfactory '
        'condition and suitable for residential habitation. Any deficiencies noted during '
        'the walk-through inspection have been documented in the Move-In Condition Report '
        'attached hereto as Exhibit A.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        'The Tenant shall use the Premises exclusively for residential purposes and shall '
        'not conduct any commercial, industrial, or professional business activities on '
        'the Premises without the prior written consent of the Landlord.'
    )
    r3.font.size = Pt(11)
    r3.font.name = 'Times New Roman'

    # --- Article III: Term ---
    doc.add_heading('Article III: Term of Lease', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'The initial term of this Lease shall commence on April 1, 2025, and shall '
        'terminate on March 31, 2026, for a total period of twelve (12) months (the '
        '"Initial Term"). Unless either party provides written notice of non-renewal at '
        'least sixty (60) days prior to the expiration of the Initial Term or any renewal '
        'term, this Lease shall automatically renew on a month-to-month basis under the '
        'same terms and conditions, except that the monthly rent may be adjusted as '
        'provided in Article IV.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'Early termination of this Lease by the Tenant shall require a minimum of ninety '
        '(90) days written notice and payment of an early termination fee equal to two (2) '
        "months' rent. The Landlord reserves the right to terminate this Lease in accordance "
        'with applicable California law upon material breach by the Tenant of any provision '
        'contained herein.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    # --- Article IV: Rent ---
    doc.add_heading('Article IV: Rent', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'The Tenant agrees to pay monthly rent in the amount of Two Thousand Eight Hundred '
        'Fifty Dollars ($2,850.00) per month, due on the first (1st) day of each calendar '
        'month during the term of this Lease. Rent shall be payable by check, electronic '
        'funds transfer, or certified bank draft made payable to Harrington Property '
        'Management LLC.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'A late fee of Seventy-Five Dollars ($75.00) shall be assessed for any rent payment '
        'received after the fifth (5th) day of the month. If rent remains unpaid for more '
        'than fifteen (15) days after the due date, an additional penalty of One Hundred '
        'Fifty Dollars ($150.00) shall be imposed. The Landlord may, at its sole discretion, '
        'increase the monthly rent by no more than five percent (5%) per annum upon sixty '
        '(60) days written notice prior to the commencement of any renewal term.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        'All rent payments shall be mailed or delivered to the Landlord at 4200 Wilshire '
        'Boulevard, Suite 310, Los Angeles, California 90010, or to such other address as '
        'the Landlord may designate in writing from time to time.'
    )
    r3.font.size = Pt(11)
    r3.font.name = 'Times New Roman'

    # --- Article V: Security Deposit ---
    doc.add_heading('Article V: Security Deposit', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'Upon execution of this Agreement, the Tenant shall deposit with the Landlord the '
        'sum of Five Thousand Seven Hundred Dollars ($5,700.00) as a security deposit (the '
        '"Security Deposit"). The Security Deposit shall be held by the Landlord in a '
        'separate interest-bearing account in accordance with California Civil Code Section '
        '1950.5 and shall be used to cover any damages to the Premises beyond normal wear '
        'and tear, unpaid rent, or other charges owed by the Tenant under this Agreement.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'Within twenty-one (21) days after the Tenant vacates the Premises, the Landlord '
        'shall provide the Tenant with an itemized statement of any deductions from the '
        'Security Deposit and shall return the remaining balance, if any, to the Tenant at '
        'the forwarding address provided by the Tenant. The Tenant may not apply the Security '
        'Deposit as payment for the last month of rent without the prior written consent of '
        'the Landlord.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    # --- Article VI: Maintenance and Repairs ---
    doc.add_heading('Article VI: Maintenance and Repairs', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'The Landlord shall be responsible for maintaining the structural components of the '
        'building, including the roof, exterior walls, foundation, plumbing systems, '
        'electrical systems, and common areas. The Landlord shall ensure that all essential '
        'services, including water, heating, and electrical supply, remain in good working '
        'order throughout the term of this Lease.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'The Tenant shall be responsible for routine maintenance and upkeep of the interior '
        'of the Premises, including but not limited to replacing light bulbs, maintaining '
        'cleanliness, preventing pest infestations, and promptly reporting any maintenance '
        'issues or needed repairs to the Landlord. The Tenant shall not make any structural '
        'alterations, additions, or improvements to the Premises without the prior written '
        'consent of the Landlord.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        'The Tenant agrees to promptly notify the Landlord of any condition that poses a '
        'health or safety hazard or that may result in damage to the Premises. Failure to '
        'report such conditions in a timely manner may result in the Tenant being held '
        'liable for any resulting damage.'
    )
    r3.font.size = Pt(11)
    r3.font.name = 'Times New Roman'

    # --- Article VII: Termination ---
    doc.add_heading('Article VII: Termination and Default', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'This Agreement may be terminated by either party upon the occurrence of any of the '
        'following events: (a) expiration of the Initial Term or any renewal term with '
        'proper notice of non-renewal; (b) mutual written agreement of both parties; '
        '(c) material breach of any provision of this Agreement by either party, provided '
        'that the non-breaching party has given written notice of the breach and the '
        'breaching party has failed to cure the breach within thirty (30) days of receipt '
        'of such notice.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'In the event of default by the Tenant, including but not limited to non-payment '
        'of rent, violation of the terms of this Agreement, or engaging in illegal activity '
        'on the Premises, the Landlord may pursue all remedies available under California '
        'law, including but not limited to unlawful detainer proceedings, recovery of unpaid '
        "rent, and recovery of the Landlord's reasonable attorney's fees and court costs."
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    # --- Article VIII: Governing Law ---
    doc.add_heading('Article VIII: Governing Law and Dispute Resolution', level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        'This Agreement shall be governed by and construed in accordance with the laws of '
        'the State of California, without regard to its conflict of laws principles. Any '
        'disputes arising out of or relating to this Agreement shall first be submitted to '
        'mediation in accordance with the rules of the American Arbitration Association. '
        'If mediation is unsuccessful, the dispute shall be resolved through binding '
        'arbitration conducted in Los Angeles County, California.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'The prevailing party in any arbitration or legal proceeding arising under this '
        "Agreement shall be entitled to recover its reasonable attorney's fees, costs, and "
        'expenses from the non-prevailing party. This Agreement constitutes the entire '
        'understanding between the parties and supersedes all prior negotiations, '
        'representations, warranties, commitments, offers, contracts, and writings of any '
        'nature, whether oral or written, with respect to the subject matter hereof.'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    # Signature block
    doc.add_paragraph()  # spacer
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(24)
    r_sig = sig.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of '
                        'the date first written above.')
    r_sig.font.size = Pt(11)
    r_sig.font.name = 'Times New Roman'
    r_sig.bold = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
