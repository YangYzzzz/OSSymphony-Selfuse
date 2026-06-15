"""
Initial Setup: Feature list document with default round bullet (U+2022)
Task ID: writer_list_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_list_002'
OUTPUT = f'{WORKDIR}/feature_list.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title heading for context
    doc.add_heading('Product Feature List', level=1)

    # Add 8 bullet list items using List Bullet style
    items = [
        'Cloud-based storage',
        'Real-time collaboration',
        'Version history tracking',
        'Role-based access control',
        'API integrations',
        'Custom dashboards',
        'Automated reporting',
        'Mobile app support',
    ]

    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Modify the numbering XML to use U+2022 (round bullet) with Arial font
    # The ListBullet style maps to abstractNum with pStyle=ListBullet
    numbering_part = doc.part.numbering_part
    if numbering_part is not None:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        numbering_xml = numbering_part._element

        # Find the abstractNum that has pStyle=ListBullet (the one used by our paragraphs)
        for abstract_num in numbering_xml.findall('w:abstractNum', ns):
            for lvl in abstract_num.findall('w:lvl', ns):
                pstyle = lvl.find('w:pStyle', ns)
                if pstyle is not None and pstyle.get(qn('w:val')) == 'ListBullet':
                    # Change bullet to U+2022 with Arial font
                    lvl_text = lvl.find('w:lvlText', ns)
                    if lvl_text is not None:
                        lvl_text.set(qn('w:val'), '\u2022')
                    # Update the run properties font to Arial (supports U+2022)
                    rpr = lvl.find('w:rPr', ns)
                    if rpr is None:
                        rpr = etree.SubElement(lvl, qn('w:rPr'))
                    rfonts = rpr.find('w:rFonts', ns)
                    if rfonts is None:
                        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
                    rfonts.set(qn('w:ascii'), 'Arial')
                    rfonts.set(qn('w:hAnsi'), 'Arial')
                    # Remove hint="default" if present (Symbol hint)
                    if rfonts.get(qn('w:hint')):
                        del rfonts.attrib[qn('w:hint')]

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
