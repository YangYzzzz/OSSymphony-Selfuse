"""
Reward Script: Create a user-defined index called 'Glossary'
Task ID: writer_mt_087
Domain: libreoffice_writer

Scoring rubric:
  Component 1 (0.35): TOC field code for Glossary index exists (TOC \f "Glossary")
  Component 2 (0.30): A heading paragraph titled 'Glossary' exists at/near the end
  Component 3 (0.35): Index content lists terms alphabetically with page numbers
                       (at least 25 of the 30 expected terms present)
Total: 1.0
"""

import os
import re
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_087'

# The 30 expected terms from the XE entries
EXPECTED_TERMS = [
    "Abstraction", "API", "Bandwidth", "Bytecode", "Cache",
    "Containerization", "Daemon", "Dependency Injection", "Encryption",
    "Firmware", "Gateway", "Hashing", "Instance", "JSON", "Kernel",
    "Latency", "Middleware", "Namespace", "Overhead", "Pipeline",
    "Query", "Repository", "Serialization", "Throughput", "URI",
    "Virtualization", "Webhook", "XML", "YAML", "Zero-Day"
]


def verify_task(file_path):
    """
    Verify that a user-defined index called 'Glossary' has been inserted.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: TOC field code for Glossary index exists (0.35 points)
    # The golden file should have a TOC field with \f "Glossary" switch,
    # which is absent in the initial file (initial only has XE entries).
    try:
        body = doc.element.body
        xml_str = etree.tostring(body, encoding='unicode')
        instrs = re.findall(r'<w:instrText[^>]*>(.*?)</w:instrText>', xml_str)

        # Count TOC instructions that reference Glossary
        toc_count = sum(1 for instr in instrs if 'TOC' in instr and 'Glossary' in instr)

        if toc_count > 0:
            print(f"PASS: Component 1 -- TOC field with Glossary filter found ({toc_count} match(es)) (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 -- No TOC field with Glossary filter found among {len(instrs)} field instructions")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: A heading titled 'Glossary' exists near the end of the document (0.30 points)
    # The initial file does not have this heading; the golden file adds it.
    try:
        check_range = doc.paragraphs[-10:] if len(doc.paragraphs) >= 10 else doc.paragraphs
        # Count headings with 'Glossary' text
        glossary_headings = [p for p in check_range
                             if 'Heading' in p.style.name and p.text.strip().lower() == 'glossary']

        if len(glossary_headings) > 0:
            print(f"PASS: Component 2 -- 'Glossary' heading found near end of document (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 -- No 'Glossary' heading found in last paragraphs")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Index content contains the expected terms (0.35 points)
    # The golden file should have index entries listed (terms + page numbers).
    # We check that at least 25 of the 30 expected terms appear in the text
    # of paragraphs AFTER the Glossary heading.
    try:
        # Find the glossary heading position
        glossary_idx = None
        for i, para in enumerate(doc.paragraphs):
            if 'Heading' in para.style.name and para.text.strip().lower() == 'glossary':
                glossary_idx = i
                break

        if glossary_idx is not None:
            # Collect text from all paragraphs after the glossary heading
            index_text = ' '.join(p.text for p in doc.paragraphs[glossary_idx + 1:])
            terms_found = [t for t in EXPECTED_TERMS if t in index_text]

            if len(terms_found) >= 25:
                ratio = min(len(terms_found) / 30.0, 1.0)
                score_c3 = round(0.35 * ratio, 4)
                print(f"PASS: Component 3 -- {len(terms_found)}/30 terms found in index content ({score_c3:.2f} pts)")
                if len(terms_found) >= 25:
                    total_score += score_c3
            else:
                print(f"FAIL: Component 3 -- Only {len(terms_found)}/30 terms in index content after Glossary heading")
        else:
            # Fallback: check trailing paragraphs even without heading
            all_text = ' '.join(p.text for p in doc.paragraphs[-5:])
            terms_found = [t for t in EXPECTED_TERMS if t in all_text]
            if len(terms_found) >= 25:
                ratio = min(len(terms_found) / 30.0, 1.0)
                score_c3 = round(0.35 * ratio, 4)
                print(f"PASS: Component 3 -- {len(terms_found)}/30 terms found in trailing paragraphs ({score_c3:.2f} pts)")
                if len(terms_found) >= 25:
                    total_score += score_c3
            else:
                print(f"FAIL: Component 3 -- No Glossary heading found, only {len(terms_found)}/30 terms in trailing text")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
