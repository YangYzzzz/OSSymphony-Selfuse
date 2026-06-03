"""
Initial Setup: Create Technical_Spec.docx with 30 user-defined index entries for 'Glossary'
Task ID: writer_mt_087
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_087'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# ---- Glossary terms with definitions ----
GLOSSARY_TERMS = [
    ("API", "Application Programming Interface, a set of protocols for building software"),
    ("Bandwidth", "Maximum rate of data transfer across a network path"),
    ("Cache", "Hardware or software component storing data for faster future requests"),
    ("Daemon", "Background process that runs without direct user interaction"),
    ("Encryption", "Process of encoding information to prevent unauthorized access"),
    ("Firmware", "Permanent software programmed into a hardware device's memory"),
    ("Gateway", "Network node that serves as an access point to another network"),
    ("Hashing", "Mapping data of arbitrary size to fixed-size values using a hash function"),
    ("Instance", "A single occurrence of an object created from a class template"),
    ("JSON", "JavaScript Object Notation, a lightweight data interchange format"),
    ("Kernel", "Core component of an operating system managing hardware resources"),
    ("Latency", "Time delay between a request and the corresponding response"),
    ("Middleware", "Software acting as a bridge between an operating system and applications"),
    ("Namespace", "Container providing scope for identifiers to avoid naming collisions"),
    ("Overhead", "Additional computation or resources required beyond the core task"),
    ("Pipeline", "Series of data processing elements connected in sequence"),
    ("Query", "Request for data or information from a database or search engine"),
    ("Repository", "Central location for storing and managing source code and artifacts"),
    ("Serialization", "Converting an object into a byte stream for storage or transmission"),
    ("Throughput", "Amount of data processed by a system in a given time period"),
    ("URI", "Uniform Resource Identifier, a string identifying a resource on a network"),
    ("Virtualization", "Creating virtual versions of hardware, storage, or network resources"),
    ("Webhook", "HTTP callback triggered by an event in a source system"),
    ("XML", "Extensible Markup Language, a markup language for encoding documents"),
    ("YAML", "Human-readable data serialization language often used for configuration"),
    ("Zero-Day", "Previously unknown vulnerability exploited before a patch is available"),
    ("Abstraction", "Hiding complex implementation details behind a simpler interface"),
    ("Bytecode", "Intermediate code compiled from source for execution by a virtual machine"),
    ("Containerization", "Packaging applications with dependencies into isolated containers"),
    ("Dependency Injection", "Design pattern where objects receive dependencies from external sources"),
]


def add_xe_field(paragraph, term, definition):
    """Insert an XE (index entry) field for a user-defined 'Glossary' index into a paragraph.

    Field code format: XE "term" \\f "Glossary" \\d "definition"
    (\\f sets the index type, but for user-defined index in LO, we use the type identifier)
    """
    # We mark the entry using a field code:  XE "term:definition" \f "Glossary"
    # The \f switch assigns the entry to a user-defined index type.

    r_begin = etree.SubElement(paragraph._element, qn('w:r'))
    fld_char_begin = etree.SubElement(r_begin, qn('w:fldChar'))
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    r_instr = etree.SubElement(paragraph._element, qn('w:r'))
    instr_text = etree.SubElement(r_instr, qn('w:instrText'))
    instr_text.set(qn('xml:space'), 'preserve')
    # Use \f "Glossary" to assign to user-defined index named Glossary
    instr_text.text = f' XE "{term}" \\f "Glossary" '

    r_sep = etree.SubElement(paragraph._element, qn('w:r'))
    fld_char_sep = etree.SubElement(r_sep, qn('w:fldChar'))
    fld_char_sep.set(qn('w:fldCharType'), 'separate')

    r_end = etree.SubElement(paragraph._element, qn('w:r'))
    fld_char_end = etree.SubElement(r_end, qn('w:fldChar'))
    fld_char_end.set(qn('w:fldCharType'), 'end')


def create_initial():
    doc = Document()

    # ---- Document styles ----
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading('Technical Specification: Cloud Infrastructure Platform', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Version info ----
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Version 3.2 | March 2026 | Confidential')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ---- Section 1: Introduction ----
    doc.add_heading('1. Introduction', level=1)
    p1 = doc.add_paragraph(
        'This document outlines the technical specifications for the Cloud Infrastructure '
        'Platform (CIP), a next-generation system designed to handle enterprise-scale workloads. '
        'The platform leverages modern architectural patterns including microservices, '
        'event-driven communication, and infrastructure-as-code principles.'
    )
    # Add XE entries inline after relevant paragraphs
    add_xe_field(p1, 'API', GLOSSARY_TERMS[0][1])

    p2 = doc.add_paragraph(
        'The CIP architecture employs a layered design where each service communicates through '
        'well-defined API endpoints. Network bandwidth allocation is dynamically managed based '
        'on real-time traffic analysis and historical throughput patterns.'
    )
    add_xe_field(p2, 'Bandwidth', GLOSSARY_TERMS[1][1])
    add_xe_field(p2, 'Throughput', GLOSSARY_TERMS[19][1])

    # ---- Section 2: System Architecture ----
    doc.add_heading('2. System Architecture', level=1)

    doc.add_heading('2.1 Caching Layer', level=2)
    p3 = doc.add_paragraph(
        'The distributed cache subsystem implements a multi-tier strategy. The L1 cache operates '
        'at the application level using in-memory stores, while the L2 cache leverages Redis '
        'clusters for cross-service data sharing. Cache invalidation follows an event-driven '
        'model to ensure consistency across all instances of the service.'
    )
    add_xe_field(p3, 'Cache', GLOSSARY_TERMS[2][1])
    add_xe_field(p3, 'Instance', GLOSSARY_TERMS[8][1])

    doc.add_heading('2.2 Background Services', level=2)
    p4 = doc.add_paragraph(
        'System daemons monitor health metrics, manage log rotation, and coordinate '
        'periodic maintenance tasks. Each daemon registers with the service mesh through '
        'the central gateway, which acts as the primary entry point for all external requests. '
        'Communication between internal services uses gRPC with Protocol Buffer serialization.'
    )
    add_xe_field(p4, 'Daemon', GLOSSARY_TERMS[3][1])
    add_xe_field(p4, 'Gateway', GLOSSARY_TERMS[6][1])
    add_xe_field(p4, 'Serialization', GLOSSARY_TERMS[18][1])

    doc.add_heading('2.3 Data Processing Pipeline', level=2)
    p5 = doc.add_paragraph(
        'The ETL pipeline processes incoming data streams through a sequence of transformation '
        'stages. Each stage in the pipeline can be independently scaled. Intermediate results '
        'are stored in a columnar format optimized for analytical queries against the data '
        'warehouse. The query engine supports both SQL and GraphQL interfaces.'
    )
    add_xe_field(p5, 'Pipeline', GLOSSARY_TERMS[15][1])
    add_xe_field(p5, 'Query', GLOSSARY_TERMS[16][1])

    # ---- Section 3: Security ----
    doc.add_heading('3. Security Framework', level=1)

    doc.add_heading('3.1 Encryption Standards', level=2)
    p6 = doc.add_paragraph(
        'All data at rest uses AES-256 encryption. Data in transit is protected by TLS 1.3. '
        'Cryptographic hashing with SHA-256 ensures message integrity across service boundaries. '
        'The security team conducts regular audits to identify potential zero-day vulnerabilities '
        'before they can be exploited in production environments.'
    )
    add_xe_field(p6, 'Encryption', GLOSSARY_TERMS[4][1])
    add_xe_field(p6, 'Hashing', GLOSSARY_TERMS[7][1])
    add_xe_field(p6, 'Zero-Day', GLOSSARY_TERMS[25][1])

    doc.add_heading('3.2 Access Control', level=2)
    p7 = doc.add_paragraph(
        'Role-based access control (RBAC) is enforced at every layer. Each namespace '
        'maintains its own set of permissions and resource quotas. Authentication tokens '
        'use JSON Web Tokens (JWT) with configurable expiration policies.'
    )
    add_xe_field(p7, 'Namespace', GLOSSARY_TERMS[13][1])
    add_xe_field(p7, 'JSON', GLOSSARY_TERMS[9][1])

    # ---- Section 4: Infrastructure ----
    doc.add_heading('4. Infrastructure Layer', level=1)

    doc.add_heading('4.1 Firmware and Hardware', level=2)
    p8 = doc.add_paragraph(
        'Server firmware is updated quarterly through an automated rollout process. '
        'The kernel configuration is optimized for high-throughput network I/O, with custom '
        'modules loaded for NVMe storage and SR-IOV network virtualization. '
        'Resource overhead is continuously monitored to identify inefficiencies.'
    )
    add_xe_field(p8, 'Firmware', GLOSSARY_TERMS[5][1])
    add_xe_field(p8, 'Kernel', GLOSSARY_TERMS[10][1])
    add_xe_field(p8, 'Virtualization', GLOSSARY_TERMS[21][1])
    add_xe_field(p8, 'Overhead', GLOSSARY_TERMS[14][1])

    doc.add_heading('4.2 Containerization Strategy', level=2)
    p9 = doc.add_paragraph(
        'All services are deployed as containers orchestrated by Kubernetes. Containerization '
        'ensures consistent behavior across development, staging, and production. Dependencies '
        'are managed through dependency injection frameworks, reducing coupling between modules. '
        'Application bytecode is compiled into optimized container images for faster startup.'
    )
    add_xe_field(p9, 'Containerization', GLOSSARY_TERMS[28][1])
    add_xe_field(p9, 'Dependency Injection', GLOSSARY_TERMS[29][1])
    add_xe_field(p9, 'Bytecode', GLOSSARY_TERMS[27][1])

    # ---- Section 5: Monitoring ----
    doc.add_heading('5. Monitoring and Observability', level=1)
    p10 = doc.add_paragraph(
        'The observability stack provides end-to-end visibility into system behavior. '
        'Latency percentiles (p50, p95, p99) are tracked for every service endpoint. '
        'Webhook integrations notify on-call engineers when alerts fire. Configuration '
        'is stored in YAML files managed through the central repository.'
    )
    add_xe_field(p10, 'Latency', GLOSSARY_TERMS[11][1])
    add_xe_field(p10, 'Webhook', GLOSSARY_TERMS[22][1])
    add_xe_field(p10, 'YAML', GLOSSARY_TERMS[24][1])
    add_xe_field(p10, 'Repository', GLOSSARY_TERMS[17][1])

    # ---- Section 6: APIs and Integration ----
    doc.add_heading('6. APIs and Integration', level=1)
    p11 = doc.add_paragraph(
        'External integrations use RESTful APIs with responses formatted in JSON or XML. '
        'Each resource is identified by a unique URI that follows the RFC 3986 standard. '
        'The middleware layer handles request transformation, authentication, and rate limiting. '
        'Abstraction layers shield downstream consumers from backend implementation changes.'
    )
    add_xe_field(p11, 'XML', GLOSSARY_TERMS[23][1])
    add_xe_field(p11, 'URI', GLOSSARY_TERMS[20][1])
    add_xe_field(p11, 'Middleware', GLOSSARY_TERMS[12][1])
    add_xe_field(p11, 'Abstraction', GLOSSARY_TERMS[26][1])

    # ---- No Glossary index inserted (that's the task!) ----

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


create_initial()
