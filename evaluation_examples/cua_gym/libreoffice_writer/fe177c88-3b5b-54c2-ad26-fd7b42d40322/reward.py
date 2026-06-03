"""
Reward Script: Format financial report summary with tab stops
Task ID: osworld_writer_tabstop_split_line_008
Domain: libreoffice_writer
Scoring:
  Component 1: RIGHT tab at 17 cm on all 15 summary lines (0.5 pts)
  Component 2: LEFT tab at 0 cm on all 15 summary lines (0.2 pts)
  Component 3: Tab character separator used in all 15 summary lines (0.3 pts)
  Total: 1.0

Task: Apply a left tab at 0 cm and right tab at 17 cm to all 15 lines in the
      financial report summary section (paragraphs 4-18), and ensure each line
      uses a tab character between account name and balance.
"""

import os
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_split_line_008'

# EMU value for 17 cm (right tab target)
RIGHT_TAB_POS_EMU = 6120130  # 17 cm in EMUs (17 * 914400 / 2.54 = 6120000, actual = 6120130)
RIGHT_TAB_TOLERANCE_EMU = 5000  # ~0.014 cm tolerance

# Summary section paragraph indices (0-based): paragraphs 4 through 18 inclusive = 15 lines
SUMMARY_PARA_START = 4
SUMMARY_PARA_END = 18  # inclusive


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: ensure the document has at least 19 paragraphs
    if len(doc.paragraphs) < 19:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 19")
        print("REWARD: 0.0")
        return 0.0

    summary_paras = doc.paragraphs[SUMMARY_PARA_START:SUMMARY_PARA_END + 1]
    num_summary = len(summary_paras)
    print(f"INFO: Checking {num_summary} summary paragraphs (indices {SUMMARY_PARA_START}-{SUMMARY_PARA_END})")

    # -------------------------------------------------------------------------
    # Component 1: RIGHT tab at 17 cm is present on all 15 summary paragraphs (0.5 pts)
    # -------------------------------------------------------------------------
    try:
        right_tab_count = 0
        for i, para in enumerate(summary_paras):
            found_right = False
            for ts in para.paragraph_format.tab_stops:
                if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                    continue
                if ts.alignment == WD_TAB_ALIGNMENT.RIGHT:
                    pos_cm = ts.position / 914400 * 2.54
                    if abs(ts.position - RIGHT_TAB_POS_EMU) <= RIGHT_TAB_TOLERANCE_EMU:
                        found_right = True
                        break
            if found_right:
                right_tab_count += 1
            else:
                print(f"  FAIL Component 1: Para {SUMMARY_PARA_START + i} missing RIGHT tab at 17 cm")

        if right_tab_count == num_summary:
            print(f"PASS: Component 1 — RIGHT tab at 17 cm found on all {num_summary} summary lines (0.5 pts)")
            total_score += 0.5
        elif right_tab_count > 0:
            partial = round(0.5 * right_tab_count / num_summary, 3)
            print(f"PARTIAL: Component 1 — RIGHT tab at 17 cm found on {right_tab_count}/{num_summary} lines ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No summary lines have RIGHT tab at 17 cm (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: LEFT tab at 0 cm is present on all 15 summary paragraphs (0.2 pts)
    # -------------------------------------------------------------------------
    try:
        left_tab_count = 0
        for i, para in enumerate(summary_paras):
            found_left = False
            for ts in para.paragraph_format.tab_stops:
                if ts.alignment == WD_TAB_ALIGNMENT.CLEAR:
                    continue
                if ts.alignment == WD_TAB_ALIGNMENT.LEFT and ts.position == 0:
                    found_left = True
                    break
            if found_left:
                left_tab_count += 1
            else:
                print(f"  FAIL Component 2: Para {SUMMARY_PARA_START + i} missing LEFT tab at 0 cm")

        if left_tab_count == num_summary:
            print(f"PASS: Component 2 — LEFT tab at 0 cm found on all {num_summary} summary lines (0.2 pts)")
            total_score += 0.2
        elif left_tab_count > 0:
            partial = round(0.2 * left_tab_count / num_summary, 3)
            print(f"PARTIAL: Component 2 — LEFT tab at 0 cm found on {left_tab_count}/{num_summary} lines ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No summary lines have LEFT tab at 0 cm (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Tab character used as separator in all 15 summary paragraphs (0.3 pts)
    # -------------------------------------------------------------------------
    try:
        tab_char_count = 0
        for i, para in enumerate(summary_paras):
            if '\t' in para.text:
                tab_char_count += 1
            else:
                print(f"  FAIL Component 3: Para {SUMMARY_PARA_START + i} has no tab character — text={repr(para.text[:60])}")

        if tab_char_count == num_summary:
            print(f"PASS: Component 3 — Tab character separator found in all {num_summary} summary lines (0.3 pts)")
            total_score += 0.3
        elif tab_char_count > 0:
            partial = round(0.3 * tab_char_count / num_summary, 3)
            print(f"PARTIAL: Component 3 — Tab character found in {tab_char_count}/{num_summary} lines ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No summary lines use tab character separator (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
