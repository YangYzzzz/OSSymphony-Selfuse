"""
Initial Setup: Financial report summary with irregular spacing (no tab stops)
Task ID: osworld_writer_tabstop_split_line_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_split_line_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Annual Financial Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Fiscal Year 2024 — Summary of Accounts')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(12)

    # Intro section
    intro = doc.add_paragraph(
        'The following report presents a consolidated summary of all active accounts '
        'and their corresponding year-end balances as reported by the Finance Department. '
        'All figures are in USD and reflect final audited values.'
    )
    intro.paragraph_format.space_after = Pt(6)

    # Section heading
    section_heading = doc.add_heading('Account Balance Summary', level=2)
    section_heading.paragraph_format.space_before = Pt(12)
    section_heading.paragraph_format.space_after = Pt(6)

    # 15 summary lines with irregular spacing (NO tab stops, irregular whitespace)
    # Account name on the left, balance on the right, separated by spaces (not tabs)
    summary_lines = [
        ('Operating Cash Account',        '     $124,850.00'),
        ('Accounts Receivable',           '       $98,342.75'),
        ('Inventory — Finished Goods',    '    $215,640.30'),
        ('Prepaid Expenses',              '         $18,920.00'),
        ('Equipment & Machinery',         '      $342,100.00'),
        ('Accumulated Depreciation',      '    -$87,450.00'),
        ('Accounts Payable',              '       -$62,310.50'),
        ('Short-Term Loans',              '         -$45,000.00'),
        ('Accrued Liabilities',           '       -$23,870.25'),
        ('Long-Term Debt',                '         -$185,000.00'),
        ('Common Stock',                  '            $500,000.00'),
        ('Retained Earnings',             '       $312,480.50'),
        ('Revenue — Product Sales',       '     $1,045,230.00'),
        ('Cost of Goods Sold',            '        -$623,145.00'),
        ('Operating Expenses',            '         -$198,760.75'),
    ]

    for account, balance in summary_lines:
        # Use irregular spaces to separate account name and balance (no tabs)
        line_text = account + balance
        para = doc.add_paragraph(line_text)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        # NO tab stops added — this is the initial state

    # Footer note
    doc.add_paragraph('')
    note = doc.add_paragraph(
        'Note: All balances are as of December 31, 2024. '
        'Figures prepared in accordance with GAAP standards. '
        'For inquiries, contact finance@company.com.'
    )
    note.paragraph_format.space_before = Pt(12)
    for run in note.runs:
        run.font.italic = True
        run.font.size = Pt(9)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
