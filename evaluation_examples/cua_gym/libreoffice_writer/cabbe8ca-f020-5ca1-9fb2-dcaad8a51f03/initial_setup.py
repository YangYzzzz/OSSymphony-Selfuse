"""
Initial Setup: Insert endnotes for cited references in a history essay
Task ID: writer_rd_055
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---- Title ----
    title = doc.add_heading('The Fall of the Roman Empire: Causes and Consequences', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Historical Analysis')
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ---- Section 1: Introduction ----
    doc.add_heading('1. Introduction', level=1)

    doc.add_paragraph(
        'The decline and fall of the Roman Empire remains one of the most debated topics '
        'in Western historiography. Spanning several centuries, the transformation of the '
        'Roman world from a unified Mediterranean empire to a collection of successor kingdoms '
        'has fascinated scholars since Edward Gibbon first published his monumental work in 1776. '
        'The causes of this decline are multifaceted and continue to generate scholarly debate.'
    )

    doc.add_paragraph(
        'The Roman Empire at its zenith controlled approximately 5 million square kilometers '
        'of territory, stretching from Britain in the northwest to Mesopotamia in the east. '
        'At its peak under Emperor Trajan in 117 AD, the empire encompassed an estimated '
        '70 million people, roughly 21 percent of the world population at the time. The '
        'administrative, military, and economic systems that sustained this vast territory '
        'were among the most sophisticated the ancient world had ever seen.'
    )

    doc.add_paragraph(
        'Yet by the late fifth century, the western half of the empire had fragmented into '
        'a series of Germanic kingdoms. The traditional date of 476 AD, when the last Western '
        'Roman Emperor Romulus Augustulus was deposed by the Germanic chieftain Odoacer, '
        'marks the conventional endpoint. However, modern historians increasingly view this '
        'as part of a longer transformation rather than a sudden collapse.'
    )

    # ---- Section 2: Economic Factors ----
    doc.add_heading('2. Economic Factors in the Decline', level=1)

    doc.add_paragraph(
        'The economic foundations of the Roman Empire began to show signs of strain as early '
        'as the third century. The Crisis of the Third Century (235-284 AD) brought widespread '
        'economic disruption, including severe inflation, disruption of trade networks, and '
        'a decline in agricultural productivity. The silver content of the denarius, which '
        'had been approximately 95 percent under Augustus, fell to less than 5 percent by '
        'the reign of Gallienus in the 260s.'
    )

    # Reference 1 inline
    doc.add_paragraph(
        'The debasement of coinage had profound effects on commerce and taxation. As Wickham '
        'has demonstrated, the fiscal machinery of the late empire became increasingly reliant '
        'on payments in kind rather than monetary transactions (see reference 1). This shift '
        'fundamentally altered the relationship between the central government and provincial '
        'populations, weakening the economic integration that had been a hallmark of Roman rule.'
    )

    doc.add_paragraph(
        'Trade networks that had connected the Mediterranean basin for centuries began to '
        'contract. Archaeological evidence from shipwreck distributions shows a marked '
        'decline in long-distance maritime commerce beginning in the late fourth century. '
        'The production of fine pottery, amphorae, and other manufactured goods declined '
        'significantly, indicating a contraction of specialized economic activity.'
    )

    doc.add_paragraph(
        'Agricultural estates, the backbone of the Roman economy, underwent significant '
        'structural changes. The villa system in provinces such as Gaul and Britain showed '
        'signs of abandonment or downsizing from the mid-fourth century onward. Large '
        'landholders consolidated their holdings while smaller farmers were increasingly '
        'bound to the land through the colonate system, a precursor to medieval serfdom.'
    )

    # ---- Section 3: Military Pressures ----
    doc.add_heading('3. Military Pressures and Barbarian Incursions', level=1)

    doc.add_paragraph(
        'The Roman military machine, once the most formidable fighting force in the ancient '
        'world, faced mounting challenges from the third century onward. The frontier system '
        'established under Augustus and refined by subsequent emperors required an enormous '
        'standing army of approximately 300,000 to 400,000 troops to maintain. The cost of '
        'this military establishment placed severe strains on the imperial treasury.'
    )

    doc.add_paragraph(
        'The composition of the Roman army itself underwent dramatic changes. By the fourth '
        'century, an increasing proportion of Roman military forces consisted of Germanic '
        'foederati, allied barbarian groups who served under their own leaders in exchange '
        'for land and subsidies. Heather argues that this barbarization of the military '
        'fundamentally compromised Roman defensive capabilities (see reference 2). The '
        'loyalties of these federate troops were often divided, and their commanders sometimes '
        'pursued agendas at odds with imperial interests.'
    )

    doc.add_paragraph(
        'The Hunnic migrations of the late fourth and fifth centuries set off a chain '
        'reaction of population movements that overwhelmed Roman frontier defenses. The '
        'crossing of the Rhine by Vandals, Alans, and Suevi on December 31, 406 AD, is '
        'often cited as a decisive moment. The subsequent sack of Rome by the Visigoths '
        'under Alaric in 410 AD sent shockwaves throughout the Roman world, though the '
        'city had already ceased to be the political capital.'
    )

    doc.add_paragraph(
        'The loss of North Africa to the Vandals in the 430s was perhaps even more '
        'strategically significant than the sack of Rome itself. North Africa had been '
        'one of the empire\'s most productive agricultural regions, supplying grain to '
        'Rome and providing substantial tax revenue. The Vandal conquest cut off these '
        'vital resources and gave the Vandals a powerful naval base from which to '
        'threaten Mediterranean shipping.'
    )

    # ---- Section 4: Political and Administrative Decline ----
    doc.add_heading('4. Political and Administrative Decline', level=1)

    doc.add_paragraph(
        'The political stability of the Roman Empire deteriorated markedly from the third '
        'century onward. Between 235 and 284 AD, at least 26 individuals claimed the '
        'imperial title, and most met violent ends. This political instability undermined '
        'effective governance, disrupted administrative continuity, and diverted resources '
        'from defense and infrastructure to civil wars.'
    )

    doc.add_paragraph(
        'Diocletian\'s reforms at the end of the third century attempted to address these '
        'problems through the creation of the Tetrarchy and a massive expansion of the '
        'bureaucracy. While these reforms temporarily stabilized the empire, they also '
        'created a vastly more expensive government apparatus. Constantine\'s foundation '
        'of Constantinople in 330 AD established a second power center that would eventually '
        'eclipse Rome in political and economic importance.'
    )

    doc.add_paragraph(
        'The division of the empire into eastern and western halves, made permanent after '
        'the death of Theodosius I in 395 AD, had uneven consequences. The Eastern Empire, '
        'with its wealthier provinces and more defensible frontiers, survived and flourished '
        'as the Byzantine Empire for nearly a millennium. The Western Empire, by contrast, '
        'faced the full brunt of barbarian pressure with fewer resources.'
    )

    doc.add_paragraph(
        'Ward-Perkins contends that the fall of the Western Empire represented a genuine '
        'catastrophe in terms of material culture, literacy, and administrative sophistication '
        '(see reference 3). Archaeological evidence from sites across Western Europe reveals '
        'dramatic declines in building quality, pottery production, and coin circulation '
        'in the fifth and sixth centuries. These material indicators suggest a significant '
        'reduction in economic complexity and living standards.'
    )

    # ---- Section 5: Religious and Cultural Transformation ----
    doc.add_heading('5. Religious and Cultural Transformation', level=1)

    doc.add_paragraph(
        'The rise of Christianity profoundly transformed Roman society and culture. From '
        'a persecuted minority religion in the first three centuries, Christianity became '
        'the official state religion under Theodosius I in 380 AD. This transformation '
        'had complex effects on the empire\'s political cohesion and cultural identity.'
    )

    doc.add_paragraph(
        'The church provided an alternative source of authority and community organization '
        'that in many areas outlasted imperial administrative structures. Bishops became '
        'important civic leaders, and monastic communities preserved literacy and learning '
        'through the upheavals of the fifth and sixth centuries. However, theological '
        'controversies, particularly the Arian dispute, also created divisions within '
        'the empire and between Roman and barbarian populations.'
    )

    doc.add_paragraph(
        'The classical educational system, built around rhetoric and the study of Greek '
        'and Latin literature, gradually gave way to a more ecclesiastically oriented '
        'curriculum. The last pagan philosophical schools were closed by Justinian in '
        '529 AD, symbolizing the end of an intellectual tradition stretching back over '
        'a thousand years. Yet the church also played a crucial role in transmitting '
        'elements of classical culture to the medieval world.'
    )

    # ---- Section 6: Conclusion ----
    doc.add_heading('6. Conclusion', level=1)

    doc.add_paragraph(
        'The fall of the Roman Empire was not a single event but a protracted process '
        'driven by the interaction of multiple factors. Economic decline, military pressure, '
        'political instability, and cultural transformation all played significant roles. '
        'No single cause can adequately explain the end of a civilization that had dominated '
        'the Mediterranean world for half a millennium.'
    )

    doc.add_paragraph(
        'Modern scholarship increasingly emphasizes the complexity of this transition and '
        'the continuities as well as the ruptures between the Roman and post-Roman worlds. '
        'While the political unity of the Western Empire was lost, many Roman institutions, '
        'legal traditions, and cultural practices survived in transformed forms. The legacy '
        'of Rome continued to shape European civilization for centuries after the last '
        'emperor was deposed.'
    )

    doc.add_paragraph(
        'Understanding the fall of Rome remains relevant not merely as an exercise in '
        'ancient history but as a case study in how complex societies respond to cumulative '
        'challenges. The interplay of economic fragility, military overextension, political '
        'dysfunction, and external pressures offers insights that resonate beyond the '
        'specific historical context of late antiquity.'
    )

    # ---- References section (plain text, no endnotes) ----
    doc.add_heading('References', level=1)

    doc.add_paragraph(
        '1. Wickham, Chris. Framing the Early Middle Ages: Europe and the Mediterranean, '
        '400-800. Oxford University Press, 2005.'
    )
    doc.add_paragraph(
        '2. Heather, Peter. The Fall of the Roman Empire: A New History of Rome and the '
        'Barbarians. Oxford University Press, 2006.'
    )
    doc.add_paragraph(
        '3. Ward-Perkins, Bryan. The Fall of Rome and the End of Civilization. Oxford '
        'University Press, 2005.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
