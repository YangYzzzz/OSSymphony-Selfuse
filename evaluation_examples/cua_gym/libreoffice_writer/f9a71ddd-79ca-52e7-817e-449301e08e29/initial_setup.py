"""
Initial Setup: Marketing Technology White Paper with inline citation markers
Task ID: writer_mktg_031
Domain: libreoffice_writer

Creates martech_whitepaper.docx with:
- 8 pages of marketing technology content
- Six inline markers [1] through [6] in body text
- A 'References' section at the end listing all 6 citations
- NO proper footnotes (agent must add these)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_031'
OUTPUT = f'{WORKDIR}/martech_whitepaper.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title page ---
    title_para = doc.add_heading('The Future of Marketing Technology: Navigating the 2025 Landscape', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('A Comprehensive Analysis of MarTech Trends, Investments, and Strategic Imperatives')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(13)

    doc.add_paragraph('')
    meta = doc.add_paragraph('Prepared by: Global Marketing Strategy Group | Q1 2025')
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'The marketing technology landscape has undergone unprecedented transformation over the past '
        'three years. Organizations that once relied on legacy CRM systems and manual campaign '
        'management are now deploying sophisticated AI-driven platforms capable of real-time '
        'personalization at scale. This white paper examines the current state of the MarTech '
        'ecosystem, identifies key investment priorities, and provides strategic guidance for '
        'marketing leaders navigating a complex and rapidly evolving environment.'
    )
    doc.add_paragraph(
        'According to recent industry research, global spending on marketing technology platforms '
        'exceeded $490 billion in 2024, representing a compound annual growth rate of 14.3% since '
        '2020. [1] The consolidation trend that characterized the early 2020s has given way to a '
        'new wave of specialized solutions, particularly in the areas of customer data platforms, '
        'conversational AI, and privacy-first analytics.'
    )
    doc.add_paragraph(
        'CMOs surveyed for this report indicated that their top three priorities for 2025 are: '
        'improving data quality and integration, demonstrating measurable ROI from technology '
        'investments, and building organizational capabilities to leverage AI effectively. These '
        'priorities reflect a maturation in how enterprises approach MarTech adoption—moving from '
        'tool acquisition to strategic capability building.'
    )

    doc.add_page_break()

    # --- Section 1: The Current MarTech Ecosystem ---
    doc.add_heading('1. The Current MarTech Ecosystem', level=1)
    doc.add_heading('1.1 Market Size and Growth Trajectory', level=2)
    doc.add_paragraph(
        'The MarTech landscape, as catalogued annually by leading industry analysts, now encompasses '
        'more than 11,000 distinct solutions spanning over 50 categories. From customer journey '
        'orchestration to programmatic advertising to marketing analytics, the breadth of available '
        'tools presents both opportunity and complexity for modern marketing organizations.'
    )
    doc.add_paragraph(
        'Enterprise adoption patterns reveal a significant bifurcation between organizations that '
        'have successfully integrated their MarTech stack and those still struggling with siloed '
        'point solutions. Forrester Research reports that companies with highly integrated stacks '
        'achieve 23% higher marketing efficiency and 31% better customer retention compared to '
        'those with fragmented toolsets. [2]'
    )
    doc.add_paragraph(
        'The dominant platforms—Adobe Experience Cloud, Salesforce Marketing Cloud, HubSpot, and '
        'Oracle CX Marketing—continue to expand their feature sets through both organic development '
        'and strategic acquisition. However, best-of-breed approaches remain popular among '
        'organizations with specialized needs or existing technical infrastructure that favors '
        'modular architectures.'
    )

    doc.add_heading('1.2 Technology Investment Patterns', level=2)
    doc.add_paragraph(
        'Analysis of marketing budget allocation data from 350 enterprise organizations reveals '
        'that technology spending now represents an average of 26.6% of total marketing budgets, '
        'up from 22.0% in 2022. This increase reflects both the growing importance of digital '
        'channels and the higher expectations placed on marketing teams to demonstrate data-driven '
        'outcomes.'
    )
    para = doc.add_paragraph(
        'The digital transformation of marketing operations has delivered measurable returns for '
        'early adopters. McKinsey analysis of companies that implemented comprehensive MarTech '
        'stacks between 2019 and 2022 found median improvements of 15-25% in marketing ROI, '
        'driven primarily by improved targeting precision and reduced waste in paid media. [3] '
        'However, these gains were not uniformly distributed—organizations that invested in change '
        'management and talent development alongside technology adoption achieved results three '
        'times better than those focused solely on tool deployment.'
    )

    doc.add_page_break()

    # --- Section 2: Artificial Intelligence in Marketing ---
    doc.add_heading('2. Artificial Intelligence in Modern Marketing Operations', level=1)
    doc.add_heading('2.1 From Automation to Intelligence', level=2)
    doc.add_paragraph(
        'The integration of artificial intelligence into marketing technology represents the most '
        'significant paradigm shift since the advent of programmatic advertising. Unlike earlier '
        'waves of marketing automation—which primarily focused on rule-based triggered campaigns '
        'and batch-and-blast email sequences—modern AI applications in marketing operate on '
        'fundamentally different principles of dynamic optimization and predictive modeling.'
    )
    doc.add_paragraph(
        'Natural language processing capabilities, now embedded in virtually every major marketing '
        'platform, have transformed how teams create content, analyze customer feedback, and '
        'interact with prospects across digital channels. Conversational AI applications—chatbots, '
        'virtual assistants, and intelligent FAQ systems—now handle an estimated 67% of initial '
        'customer service interactions at companies that have deployed these solutions, according '
        'to the HubSpot State of Inbound report. [4]'
    )

    doc.add_heading('2.2 Predictive Analytics and Customer Intelligence', level=2)
    doc.add_paragraph(
        'Customer data platforms have emerged as the foundational infrastructure layer for '
        'AI-driven marketing. By unifying behavioral data, transaction history, demographic '
        'information, and real-time engagement signals, CDPs enable the sophisticated modeling '
        'required for next-generation personalization at scale.'
    )
    doc.add_paragraph(
        'Predictive lead scoring, propensity modeling, and churn prediction algorithms have '
        'matured significantly over the past two years. Organizations deploying these capabilities '
        'report meaningful improvements in sales efficiency and customer lifetime value. The '
        'challenge, however, lies in model governance—ensuring that algorithmic recommendations '
        'remain aligned with business objectives and do not inadvertently introduce bias or '
        'violate privacy regulations.'
    )

    doc.add_page_break()

    # --- Section 3: Privacy, Data Governance ---
    doc.add_heading('3. Privacy-First Marketing and Data Governance', level=1)
    doc.add_heading('3.1 The Post-Cookie Paradigm', level=2)
    doc.add_paragraph(
        'The deprecation of third-party cookies represents a fundamental restructuring of the '
        'digital advertising ecosystem. While Google\'s repeated delays in removing cookies from '
        'Chrome have extended the timeline for full transition, forward-thinking organizations '
        'have accelerated their first-party data strategies regardless of the regulatory calendar.'
    )
    doc.add_paragraph(
        'IDC research indicates that worldwide spending on marketing data management and privacy '
        'compliance solutions will reach $8.7 billion by the end of 2025, growing at a CAGR of '
        '19.2% through 2027. [5] This investment reflects the strategic priority organizations '
        'place on building sustainable, consent-based relationships with their audiences—a '
        'necessity in markets subject to GDPR, CCPA, and increasingly stringent data protection '
        'regulations globally.'
    )

    doc.add_heading('3.2 Consent Management and First-Party Data Strategy', level=2)
    doc.add_paragraph(
        'Leading organizations are reimagining their value exchange with customers, offering '
        'genuine utility in return for data sharing consent. Loyalty programs, personalized '
        'content hubs, interactive tools, and exclusive community access have emerged as effective '
        'mechanisms for building rich first-party data assets while delivering customer value.'
    )
    doc.add_paragraph(
        'The technical infrastructure required to support sophisticated consent management—'
        'including preference centers, data subject access request workflows, and automated '
        'compliance reporting—has become a standard component of enterprise MarTech stacks. '
        'Vendors that can demonstrate robust privacy-by-design architectures have gained '
        'significant competitive advantage in enterprise procurement processes.'
    )

    doc.add_page_break()

    # --- Section 4: The CMO Perspective ---
    doc.add_heading('4. Strategic Priorities: The CMO Perspective', level=1)
    doc.add_heading('4.1 Balancing Innovation and Efficiency', level=2)
    doc.add_paragraph(
        'Chief Marketing Officers face increasing pressure to demonstrate tangible business impact '
        'from technology investments while simultaneously managing costs and organizational '
        'complexity. The proliferation of point solutions that characterized MarTech expansion '
        'in the 2015-2020 period has left many organizations with bloated stacks—high maintenance '
        'costs, integration challenges, and fragmented data landscapes that undermine the very '
        'efficiency gains the tools were meant to deliver.'
    )
    doc.add_paragraph(
        'Deloitte\'s CMO Survey for Q4 2025 found that 61% of marketing leaders plan to consolidate '
        'or rationalize their technology stack over the next 18 months, down from 78% who expressed '
        'similar intentions in 2022. [6] This moderation reflects a more nuanced approach to '
        'rationalization—one that recognizes the value of specialized tools in high-impact use cases '
        'while still driving toward greater integration and operational simplicity.'
    )

    doc.add_heading('4.2 Building MarTech Competency', level=2)
    doc.add_paragraph(
        'Technology adoption without corresponding organizational capability development is widely '
        'recognized as a primary driver of MarTech failure. Research consistently shows that the '
        'human element—user adoption, process integration, data literacy, and change management—'
        'accounts for the majority of variance in MarTech outcomes.'
    )
    doc.add_paragraph(
        'Progressive organizations are investing in dedicated MarTech operations teams, sometimes '
        'called "marketing operations" or "revenue operations," to manage technology governance, '
        'vendor relationships, and continuous optimization. These teams serve as the bridge between '
        'marketing strategy and technical implementation, ensuring that the stack evolves in '
        'alignment with business priorities rather than technology trends.'
    )

    doc.add_page_break()

    # --- Section 5: Implementation Recommendations ---
    doc.add_heading('5. Implementation Recommendations', level=1)
    doc.add_paragraph(
        'Based on the research and analysis presented in this white paper, we offer the following '
        'strategic recommendations for marketing leaders seeking to maximize the value of their '
        'MarTech investments:'
    )
    doc.add_paragraph('Prioritize integration over feature breadth.', style='List Number')
    doc.add_paragraph(
        'When evaluating new platforms, weight integration capabilities and API ecosystem quality '
        'as heavily as feature functionality. A tool that works seamlessly with your existing stack '
        'will consistently outperform a more feature-rich solution that creates data silos.'
    )
    doc.add_paragraph('Invest in data quality infrastructure before advanced applications.', style='List Number')
    doc.add_paragraph(
        'AI and advanced analytics capabilities are only as effective as the underlying data they '
        'operate on. Establish robust data governance, deduplication, and enrichment processes '
        'before deploying sophisticated modeling applications.'
    )
    doc.add_paragraph('Develop a clear privacy-first data strategy.', style='List Number')
    doc.add_paragraph(
        'Build consent management and first-party data acquisition into your foundational '
        'architecture rather than treating them as compliance add-ons. Organizations that lead '
        'in this area will have sustainable competitive advantages as third-party data sources '
        'continue to erode.'
    )
    doc.add_paragraph('Measure technology ROI with rigor.', style='List Number')
    doc.add_paragraph(
        'Establish clear success metrics for each component of your MarTech stack before '
        'implementation. Regular performance reviews against these metrics enable evidence-based '
        'decisions about investment continuation, expansion, or consolidation.'
    )

    doc.add_page_break()

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The marketing technology landscape of 2025 rewards organizations that approach technology '
        'as a strategic enabler rather than a tactical solution. The companies achieving the '
        'greatest returns from their MarTech investments share common characteristics: strong '
        'executive alignment on technology strategy, rigorous data governance practices, sustained '
        'investment in organizational capability, and a relentless focus on measuring and '
        'demonstrating business impact.'
    )
    doc.add_paragraph(
        'As AI capabilities continue to mature and privacy regulations continue to evolve, the '
        'marketing technology ecosystem will undergo further transformation. Organizations that '
        'build flexible, integration-ready architectures today will be best positioned to adopt '
        'the next generation of capabilities—whatever form they ultimately take. The imperative '
        'for marketing leaders is not to predict the future of MarTech, but to build the '
        'organizational foundations that enable rapid adaptation to whatever that future holds.'
    )

    doc.add_page_break()

    # --- References section (this will be removed in golden state) ---
    doc.add_heading('References', level=1)
    ref_lines = [
        "[1] Gartner, 'Marketing Technology Survey 2025'",
        "[2] Forrester, 'The State of MarTech'",
        "[3] McKinsey, 'Digital Marketing ROI Analysis'",
        "[4] HubSpot, 'State of Inbound 2025'",
        "[5] IDC, 'Worldwide MarTech Spending Guide'",
        "[6] Deloitte, 'CMO Survey Q4 2025'",
    ]
    for ref in ref_lines:
        doc.add_paragraph(ref)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
