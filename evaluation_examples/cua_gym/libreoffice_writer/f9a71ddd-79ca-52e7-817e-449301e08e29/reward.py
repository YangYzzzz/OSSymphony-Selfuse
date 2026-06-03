"""
Reward Script: Convert inline [1]-[6] markers to proper Writer footnotes
Task ID: writer_mktg_031
Domain: libreoffice_writer
Scoring:
  Component 1: footnotes.xml exists with 6 footnotes and correct citation text  (0.30 pts)
  Component 2: 6 footnoteReference elements in document body (markers replaced)  (0.30 pts)
  Component 3: No inline [N] markers remain in body paragraphs                   (0.20 pts)
  Component 4: 'References' section removed from document end                    (0.20 pts)
  Total: 1.0
"""

import os
import zipfile
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_031'

# Expected footnote texts (from task context ground truth, without leading index)
EXPECTED_FOOTNOTES = {
    1: "Gartner, 'Marketing Technology Survey 2025'",
    2: "Forrester, 'The State of MarTech'",
    3: "McKinsey, 'Digital Marketing ROI Analysis'",
    4: "HubSpot, 'State of Inbound 2025'",
    5: "IDC, 'Worldwide MarTech Spending Guide'",
    6: "Deloitte, 'CMO Survey Q4 2025'",
}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must be readable as a valid docx (zip)
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            zip_names = z.namelist()
            has_footnotes_xml = 'word/footnotes.xml' in zip_names
            if 'word/document.xml' not in zip_names:
                print("CRITICAL: word/document.xml not found in docx — corrupt file")
                print("REWARD: 0.0")
                return 0.0
            doc_xml = z.read('word/document.xml').decode('utf-8')
            if has_footnotes_xml:
                footnotes_xml = z.read('word/footnotes.xml').decode('utf-8')
            else:
                footnotes_xml = ""
    except Exception as e:
        print(f"CRITICAL: Cannot read file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: footnotes.xml exists with 6 proper footnotes + correct text
    # This FAILS on initial (no footnotes.xml) and PASSES on golden (6 footnotes)
    # -------------------------------------------------------------------------
    try:
        if not has_footnotes_xml:
            print("FAIL: Component 1 — word/footnotes.xml does not exist in the docx")
        else:
            # Parse footnote IDs and their text content
            # Pattern: <w:footnote w:id="N"> ... <w:t>text</w:t> ...
            # Skip id="-1" (separator) and id="0" (continuationSeparator)
            ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

            import lxml.etree as etree
            root = etree.fromstring(footnotes_xml.encode('utf-8'))

            found_footnotes = {}
            for fn in root.findall(f'{{{ns_w}}}footnote'):
                fn_id_str = fn.get(f'{{{ns_w}}}id')
                fn_type = fn.get(f'{{{ns_w}}}type')
                # Skip separator footnotes (id=-1, 0 or type=separator/continuationSeparator)
                if fn_type in ('separator', 'continuationSeparator'):
                    continue
                try:
                    fn_id = int(fn_id_str)
                except (TypeError, ValueError):
                    continue
                if fn_id < 1:
                    continue
                # Collect all text within the footnote
                texts = []
                for t in fn.findall(f'.//{{{ns_w}}}t'):
                    if t.text:
                        texts.append(t.text)
                full_text = ''.join(texts).strip()
                found_footnotes[fn_id] = full_text

            num_found = len(found_footnotes)
            print(f"INFO: Found {num_found} content footnotes (IDs: {sorted(found_footnotes.keys())})")

            if num_found < 6:
                print(f"FAIL: Component 1 — expected 6 footnotes, found {num_found}")
            else:
                # Verify text content for each expected footnote
                correct_count = 0
                for fn_id, expected_text in EXPECTED_FOOTNOTES.items():
                    if fn_id in found_footnotes:
                        actual_text = found_footnotes[fn_id]
                        # Use case-insensitive substring match to be tolerant of minor spacing
                        if expected_text.lower() in actual_text.lower() or actual_text.lower() in expected_text.lower():
                            correct_count += 1
                            print(f"PASS: Footnote {fn_id} text matches: {actual_text!r}")
                        else:
                            print(f"FAIL: Footnote {fn_id} text mismatch — expected contains {expected_text!r}, got {actual_text!r}")
                    else:
                        print(f"FAIL: Footnote {fn_id} missing")

                if correct_count == 6:
                    print(f"PASS: Component 1 — all 6 footnotes present with correct citation text (0.3 pts)")
                    total_score += 0.3
                elif correct_count >= 4:
                    print(f"PARTIAL: Component 1 — {correct_count}/6 footnotes correct (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 1 — only {correct_count}/6 footnotes have correct text")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")
        import traceback
        traceback.print_exc()

    # -------------------------------------------------------------------------
    # Component 2: 6 footnoteReference elements in document body
    # This FAILS on initial (no footnoteReference elements) and PASSES on golden
    # -------------------------------------------------------------------------
    try:
        # Find all <w:footnoteReference w:id="N"/> in document.xml
        footnote_refs = re.findall(
            r'<w:footnoteReference\s[^>]*w:id=["\'](\d+)["\'][^/]*/>', doc_xml
        )
        # Also try alternate attribute order
        footnote_refs2 = re.findall(
            r'<w:footnoteReference[^>]*/>', doc_xml
        )
        num_refs = len(footnote_refs2)
        print(f"INFO: Found {num_refs} footnoteReference elements in document body")

        if num_refs >= 6:
            print(f"PASS: Component 2 — {num_refs} footnoteReference elements present (0.3 pts)")
            total_score += 0.3
        elif num_refs > 0:
            partial = round(0.3 * num_refs / 6, 2)
            print(f"PARTIAL: Component 2 — only {num_refs}/6 footnoteReference elements ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — no footnoteReference elements found in body (inline markers not converted)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: No inline [N] markers remain in paragraph body text
    # This FAILS on initial (markers present) and PASSES on golden (markers removed)
    # -------------------------------------------------------------------------
    try:
        from docx import Document
        doc = Document(file_path)

        inline_markers_found = []
        # Only check body paragraphs (not footnotes themselves)
        for para in doc.paragraphs:
            text = para.text
            for n in range(1, 7):
                marker = f'[{n}]'
                if marker in text:
                    inline_markers_found.append((n, text[:80]))

        if not inline_markers_found:
            print("PASS: Component 3 — no inline [1]-[6] markers remain in body paragraphs (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — {len(inline_markers_found)} inline markers still present:")
            for n, ctx in inline_markers_found:
                print(f"  [{n}] found in: {ctx!r}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: 'References' section removed from document
    # This FAILS on initial (References heading + 6 entries exist) and PASSES on golden
    # -------------------------------------------------------------------------
    try:
        from docx import Document
        doc = Document(file_path)

        # Count 'References' heading paragraphs and reference list entries
        references_heading_count = 0
        reference_entry_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else ''
            # Check for 'References' heading
            if text.lower() == 'references' and 'heading' in style.lower():
                references_heading_count += 1
                print(f"FAIL: Found 'References' heading paragraph (style={style!r})")
            # Check for reference list entries like "[1] Gartner..."
            if re.match(r'^\[([1-6])\]\s+\w', text):
                reference_entry_count += 1
                print(f"FAIL: Found reference list entry: {text!r}")

        if references_heading_count == 0 and reference_entry_count == 0:
            print("PASS: Component 4 — 'References' section removed from document (0.2 pts)")
            total_score += 0.2
        else:
            if references_heading_count > 0:
                print(f"FAIL: Component 4 — found {references_heading_count} 'References' heading(s) still present")
            if reference_entry_count > 0:
                print(f"FAIL: Component 4 — found {reference_entry_count} reference list entries still present")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/martech_whitepaper.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
