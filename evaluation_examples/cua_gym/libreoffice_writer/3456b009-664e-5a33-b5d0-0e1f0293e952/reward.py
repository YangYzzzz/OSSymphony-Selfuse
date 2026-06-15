"""
Reward Script: Change line spacing of second paragraph from double to single
Task ID: wrpara_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): line_spacing value is 1.0 (single) on paragraph 5
  Component 2 (0.3): line_spacing_rule is SINGLE on paragraph 5
  Component 3 (0.2): paragraph text content is preserved after spacing change
"""

import os

from docx import Document
from docx.enum.text import WD_LINE_SPACING

WORKDIR = '/home/user'
TASK_ID = 'wrpara_004'

# The paragraph that should change: index 5 in the document
# (the memo body paragraph starting with "As part of our ongoing effort...")
TARGET_PARA_INDEX = 5

# Expected text prefix to confirm we have the right paragraph
EXPECTED_TEXT_PREFIX = "As part of our ongoing effort"


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has enough paragraphs
    if len(doc.paragraphs) < TARGET_PARA_INDEX + 1:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least {TARGET_PARA_INDEX + 1}")
        print("REWARD: 0.0")
        return 0.0

    target_para = doc.paragraphs[TARGET_PARA_INDEX]

    # Precondition: confirm we have the right paragraph by checking text prefix
    if not target_para.text.startswith(EXPECTED_TEXT_PREFIX):
        print(f"CRITICAL: Paragraph {TARGET_PARA_INDEX} does not start with expected text.")
        print(f"  Found: {repr(target_para.text[:60])}")
        print("REWARD: 0.0")
        return 0.0

    pf = target_para.paragraph_format

    # Component 1: line_spacing value is 1.0 (single spacing) (0.5 points)
    # In initial_env this is 2.0 (double). In golden_env it should be 1.0.
    try:
        ls = pf.line_spacing
        if ls is not None and float(ls) == 1.0:
            print(f"PASS: Component 1 - line_spacing is {ls} (single) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 - Expected line_spacing=1.0, found: {ls}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: line_spacing_rule is SINGLE (0.3 points)
    # In initial_env this is DOUBLE (2). In golden_env it should be SINGLE (0).
    try:
        lsr = pf.line_spacing_rule
        if lsr == WD_LINE_SPACING.SINGLE:
            print(f"PASS: Component 2 - line_spacing_rule is SINGLE (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 - Expected line_spacing_rule=SINGLE, found: {lsr}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Paragraph text is preserved (not corrupted) (0.2 points)
    # The text should still start with the expected prefix and have substantial content.
    # This check only awards points when COMBINED with the spacing change (Component 1 must pass).
    # This ensures it doesn't score on initial_env where spacing is still double.
    try:
        text = target_para.text.strip()
        text_intact = (
            text.startswith(EXPECTED_TEXT_PREFIX)
            and len(text) > 100  # paragraph has substantial content
        )
        # Only award if the spacing change was also made (anchored to task change)
        if text_intact and total_score >= 0.5:
            print(f"PASS: Component 3 - Text preserved after spacing change (0.2 pts)")
            total_score += 0.2
        elif not text_intact:
            print(f"FAIL: Component 3 - Text appears corrupted or truncated (len={len(text)})")
        else:
            print(f"FAIL: Component 3 - Spacing not changed, so text preservation check skipped")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
