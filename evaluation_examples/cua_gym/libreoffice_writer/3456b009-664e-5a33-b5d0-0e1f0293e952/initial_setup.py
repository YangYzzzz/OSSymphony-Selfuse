"""
Initial Setup: Change the line spacing of the second paragraph from double spacing to single spacing.
Task ID: wrpara_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Paragraph 1: Subject Line (single-spaced) ---
    p1 = doc.add_paragraph()
    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_after = Pt(6)
    run1 = p1.add_run('MEMORANDUM')
    run1.bold = True
    run1.font.size = Pt(16)
    run1.font.name = 'Calibri'

    p1a = doc.add_paragraph()
    p1a.paragraph_format.line_spacing = 1.0
    p1a.paragraph_format.space_after = Pt(2)
    r_to = p1a.add_run('To: ')
    r_to.bold = True
    r_to.font.name = 'Calibri'
    r_to.font.size = Pt(11)
    r_to_val = p1a.add_run('All Department Managers')
    r_to_val.font.name = 'Calibri'
    r_to_val.font.size = Pt(11)

    p1b = doc.add_paragraph()
    p1b.paragraph_format.line_spacing = 1.0
    p1b.paragraph_format.space_after = Pt(2)
    r_from = p1b.add_run('From: ')
    r_from.bold = True
    r_from.font.name = 'Calibri'
    r_from.font.size = Pt(11)
    r_from_val = p1b.add_run('Rebecca Torres, VP of Operations')
    r_from_val.font.name = 'Calibri'
    r_from_val.font.size = Pt(11)

    p1c = doc.add_paragraph()
    p1c.paragraph_format.line_spacing = 1.0
    p1c.paragraph_format.space_after = Pt(2)
    r_date = p1c.add_run('Date: ')
    r_date.bold = True
    r_date.font.name = 'Calibri'
    r_date.font.size = Pt(11)
    r_date_val = p1c.add_run('March 28, 2026')
    r_date_val.font.name = 'Calibri'
    r_date_val.font.size = Pt(11)

    p1d = doc.add_paragraph()
    p1d.paragraph_format.line_spacing = 1.0
    p1d.paragraph_format.space_after = Pt(12)
    r_subj = p1d.add_run('Subject: ')
    r_subj.bold = True
    r_subj.font.name = 'Calibri'
    r_subj.font.size = Pt(11)
    r_subj_val = p1d.add_run('Updated Procedures for Quarterly Budget Reviews')
    r_subj_val.font.name = 'Calibri'
    r_subj_val.font.size = Pt(11)

    # --- Paragraph 2: Memo body (DOUBLE-spaced — agent must change to single) ---
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 2.0
    p2.paragraph_format.space_after = Pt(12)
    body_text = (
        'As part of our ongoing effort to streamline financial reporting across all divisions, '
        'the Finance Committee has approved a revised set of procedures for conducting quarterly '
        'budget reviews. Effective April 15, 2026, each department will be required to submit a '
        'preliminary budget summary no later than the first Friday of each quarter. These summaries '
        'should include year-to-date expenditures, projected costs for the upcoming quarter, and any '
        'variance explanations exceeding five percent of the approved annual budget. Please ensure '
        'that your teams are familiar with the updated submission template, which will be distributed '
        'via the internal portal by April 1, 2026.'
    )
    r2 = p2.add_run(body_text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)

    # --- Paragraph 3: Single-spaced ---
    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1.0
    p3.paragraph_format.space_after = Pt(12)
    body3 = (
        'The revised process is designed to reduce redundancies in data collection and ensure that '
        'leadership has timely access to accurate financial information. Department liaisons will '
        'receive training on the new template during the week of April 7 through April 11. If you '
        'have questions or require additional support, please contact the Finance Office at '
        'extension 4420 or email finance-reviews@company.com.'
    )
    r3 = p3.add_run(body3)
    r3.font.name = 'Calibri'
    r3.font.size = Pt(11)

    # --- Paragraph 4: Single-spaced ---
    p4 = doc.add_paragraph()
    p4.paragraph_format.line_spacing = 1.0
    p4.paragraph_format.space_after = Pt(6)
    body4 = (
        'Thank you for your cooperation in making this transition as smooth as possible. Your '
        'commitment to accurate and timely reporting is essential to the continued success of our '
        'organization. I look forward to reviewing the first round of submissions under the new '
        'guidelines and am confident this process will serve us well going forward.'
    )
    r4 = p4.add_run(body4)
    r4.font.name = 'Calibri'
    r4.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
