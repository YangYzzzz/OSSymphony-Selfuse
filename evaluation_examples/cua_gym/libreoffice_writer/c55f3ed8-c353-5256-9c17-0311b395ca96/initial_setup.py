"""
Initial Setup: Create steps_doc.docx with heading and body text (no shapes)
Task ID: writer_obj_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_064'
OUTPUT = f'{WORKDIR}/steps_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Heading: Getting Started ---
    heading = doc.add_heading('Getting Started', level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    # --- Introductory paragraph ---
    intro = doc.add_paragraph(
        'This guide walks you through the essential steps to get started with the platform. '
        'Follow the steps below to set up your environment and begin working effectively.'
    )
    intro.paragraph_format.space_after = Pt(8)

    # --- Body content ---
    doc.add_paragraph(
        'Before you begin, make sure you have the necessary permissions and access credentials. '
        'Contact your system administrator if you need assistance with account setup.'
    )

    doc.add_paragraph(
        'The onboarding process is designed to be straightforward. Each step builds on the previous one, '
        'so it is important to complete them in order. Take your time to read through each section carefully.'
    )

    doc.add_paragraph(
        'If you encounter any issues during the setup process, refer to the troubleshooting section at the end '
        'of this document or reach out to the support team at support@example.com.'
    )

    # --- Section divider paragraph (placeholder for where shapes will go) ---
    placeholder = doc.add_paragraph('')
    placeholder.paragraph_format.space_before = Pt(12)
    placeholder.paragraph_format.space_after = Pt(12)

    # --- Additional content ---
    doc.add_paragraph(
        'Once you have completed all the steps, you will have a fully configured workspace ready for use. '
        'Remember to save your settings and log out properly at the end of each session.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
