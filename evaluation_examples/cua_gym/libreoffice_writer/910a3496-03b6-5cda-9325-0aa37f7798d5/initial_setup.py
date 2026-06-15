"""
Initial Setup: Apply 'Strong Emphasis' character style to all 'MeridianPro' instances
Task ID: writer_biz_076
Domain: libreoffice_writer

Creates a business document where 'MeridianPro' appears 12 times with inconsistent
formatting (some bold, some regular, some italic) to simulate a real-world branding
consistency problem.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_text_with_product(para, text, product_bold=None, product_italic=None):
    """
    Add text to paragraph, splitting around 'MeridianPro' so product name
    gets its own run with specified formatting.
    Returns count of MeridianPro instances added.
    """
    parts = text.split('MeridianPro')
    count = 0
    for i, part in enumerate(parts):
        if part:
            para.add_run(part)
        if i < len(parts) - 1:
            run = para.add_run('MeridianPro')
            if product_bold is not None:
                run.bold = product_bold
            if product_italic is not None:
                run.italic = product_italic
            count += 1
    return count


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('MeridianPro Product Strategy Report', level=0)
    # The heading itself contains MeridianPro (instance 1) - in heading style, not Strong Emphasis
    # We need to control the run. Headings auto-create a run. Let's manually build it.
    # Clear auto-generated runs and rebuild
    for run in title.runs:
        run.clear()
    title.clear()
    run1 = title.add_run('MeridianPro')  # instance 1 - bold (inherits from heading style)
    run2 = title.add_run(' Product Strategy Report')

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)

    p1 = doc.add_paragraph()
    p1.add_run('This report outlines the strategic direction for ')
    r = p1.add_run('MeridianPro')  # instance 2 - regular (no special formatting)
    p1.add_run(' over the next fiscal year. Our flagship enterprise resource planning solution '
               'has shown consistent growth across all market segments since its launch in Q3 2024.')

    p2 = doc.add_paragraph()
    p2.add_run('Following extensive market research conducted by the Analytics Division, the ')
    r = p2.add_run('MeridianPro')  # instance 3 - italic
    r.italic = True
    p2.add_run(' platform is positioned to capture an additional 15% market share in the mid-enterprise '
               'segment by the end of 2026.')

    # --- Market Analysis ---
    doc.add_heading('Market Analysis', level=1)

    p3 = doc.add_paragraph()
    p3.add_run('The competitive landscape for enterprise resource planning has shifted significantly. ')
    r = p3.add_run('MeridianPro')  # instance 4 - bold
    r.bold = True
    p3.add_run(' differentiates itself through its modular architecture, allowing clients to deploy '
               'only the components they need. Unlike competing solutions from Terraform Systems and '
               'Apex Digital, ')
    r = p3.add_run('MeridianPro')  # instance 5 - regular
    p3.add_run(' offers seamless integration with legacy infrastructure.')

    p4 = doc.add_paragraph()
    p4.add_run('Our recent customer satisfaction survey revealed that 87% of ')
    r = p4.add_run('MeridianPro')  # instance 6 - italic
    r.italic = True
    p4.add_run(' users rated the platform as "essential" or "very important" to their daily operations. '
               'This represents a 12-point increase from the previous year\'s results.')

    # --- Product Roadmap ---
    doc.add_heading('Product Roadmap', level=1)

    p5 = doc.add_paragraph()
    p5.add_run('The development team has outlined three major milestones for ')
    r = p5.add_run('MeridianPro')  # instance 7 - bold
    r.bold = True
    p5.add_run(' version 4.0:')

    # Bullet points
    b1 = doc.add_paragraph(style='List Bullet')
    b1.add_run('AI-Powered Analytics Module: Leverage machine learning to provide predictive '
               'insights directly within the ')
    r = b1.add_run('MeridianPro')  # instance 8 - regular
    b1.add_run(' dashboard.')

    b2 = doc.add_paragraph(style='List Bullet')
    b2.add_run('Enhanced Security Framework: Implement zero-trust architecture across all ')
    r = b2.add_run('MeridianPro')  # instance 9 - italic
    r.italic = True
    b2.add_run(' endpoints and API connections.')

    b3 = doc.add_paragraph(style='List Bullet')
    b3.add_run('Cross-Platform Mobile Client: Extend ')
    r = b3.add_run('MeridianPro')  # instance 10 - bold
    r.bold = True
    b3.add_run(' functionality to iOS and Android devices with offline capability.')

    # --- Financial Projections ---
    doc.add_heading('Financial Projections', level=1)

    p6 = doc.add_paragraph()
    p6.add_run('Based on current subscription growth trends, ')
    r = p6.add_run('MeridianPro')  # instance 11 - regular
    p6.add_run(' is projected to generate $45.2 million in annual recurring revenue by Q4 2026. '
               'The Enterprise tier accounts for 62% of total revenue, while the Professional '
               'and Starter tiers contribute 28% and 10% respectively.')

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)

    p7 = doc.add_paragraph()
    p7.add_run('The strategic investments outlined in this report position ')
    r = p7.add_run('MeridianPro')  # instance 12 - italic
    r.italic = True
    p7.add_run(' for sustained growth in an increasingly competitive market. The combination of '
               'innovative product development, targeted marketing campaigns, and a customer-centric '
               'approach to service delivery will ensure that our platform continues to lead the '
               'enterprise resource planning space through 2027 and beyond.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count MeridianPro instances for verification
    doc_check = Document(OUTPUT)
    count = 0
    for para in doc_check.paragraphs:
        count += para.text.count('MeridianPro')
    print(f'MeridianPro instances found: {count}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
