"""
Initial Setup: Engineering Report with wide table on page 5
Task ID: writer_af_032
Domain: libreoffice_writer

Creates a 10-page Engineering Report document in portrait orientation.
Page 5 contains a wide table that is cut off on the right side.
All pages are portrait. No landscape sections or custom page styles exist.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_af_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body_paragraph(doc, text):
    """Add a normal paragraph with body text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def create_initial():
    doc = Document()

    # Set default page to portrait A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ====== PAGE 1: Title Page ======
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("Meridian Systems Engineering Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Quarterly Infrastructure Assessment — Q1 2025")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph("")
    meta = doc.add_paragraph("Prepared by: Dr. Elena Vasquez, Lead Systems Architect")
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta2 = doc.add_paragraph("Reviewed by: James Thornton, VP of Engineering")
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta3 = doc.add_paragraph("Date: March 28, 2025")
    meta3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta4 = doc.add_paragraph("Document Classification: Internal — Confidential")
    meta4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for p in [meta, meta2, meta3, meta4]:
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ====== PAGE 2: Table of Contents / Executive Summary ======
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary .......................... 2",
        "2. System Architecture Overview ............... 3",
        "3. Performance Metrics ........................ 4",
        "4. Infrastructure Capacity Planning ........... 5",
        "5. Security Audit Results ..................... 6",
        "6. Deployment Timeline ........................ 7",
        "7. Budget Allocation .......................... 8",
        "8. Risk Assessment ............................ 9",
        "9. Recommendations & Next Steps ............... 10",
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_paragraph("")
    add_heading_styled(doc, "1. Executive Summary", level=1)
    add_body_paragraph(doc, (
        "This report provides a comprehensive assessment of Meridian Systems' infrastructure "
        "for Q1 2025. Over the past quarter, our engineering team has completed significant "
        "upgrades to the core processing pipeline, achieving a 34% improvement in throughput "
        "and a 12% reduction in operational costs. The migration to the new Kubernetes-based "
        "orchestration layer is 78% complete, with full deployment targeted for April 15, 2025."
    ))
    add_body_paragraph(doc, (
        "Key highlights include the successful deployment of the Aurora data processing cluster, "
        "which now handles 2.4 million transactions per hour, the completion of Phase 2 security "
        "hardening across all production environments, and the onboarding of three new regional "
        "data centers in Frankfurt, Singapore, and Sao Paulo."
    ))

    doc.add_page_break()

    # ====== PAGE 3: System Architecture ======
    add_heading_styled(doc, "2. System Architecture Overview", level=1)
    add_body_paragraph(doc, (
        "The Meridian platform operates on a microservices architecture deployed across "
        "five availability zones. The primary compute layer consists of 847 container instances "
        "managed by our custom orchestration framework, CodexDeploy v3.2. Each microservice "
        "communicates via gRPC with a fallback REST API layer for external integrations."
    ))
    add_heading_styled(doc, "2.1 Core Components", level=2)
    components = [
        ("API Gateway (Envoy Proxy)", "Routes 15M+ requests/day with <2ms p99 latency"),
        ("Authentication Service", "OAuth 2.0 + SAML 2.0 with hardware token support"),
        ("Data Processing Pipeline", "Apache Kafka + custom stream processors"),
        ("Storage Layer", "PostgreSQL 16 (primary) + Redis 7.2 (cache) + S3 (blob)"),
        ("Monitoring Stack", "Prometheus + Grafana + custom alerting via PagerDuty"),
    ]
    for name, desc in components:
        p = doc.add_paragraph(style="List Bullet")
        run_bold = p.add_run(f"{name}: ")
        run_bold.bold = True
        p.add_run(desc)

    add_heading_styled(doc, "2.2 Network Topology", level=2)
    add_body_paragraph(doc, (
        "Inter-service communication uses a service mesh based on Istio 1.20, providing "
        "mutual TLS encryption, traffic management, and observability. The edge network "
        "leverages Cloudflare Enterprise for DDoS mitigation and global load balancing, "
        "achieving 99.997% uptime in Q1."
    ))
    add_body_paragraph(doc, (
        "The backbone network operates at 100 Gbps between primary data centers with automatic "
        "failover to secondary links. Regional data centers connect via dedicated fiber with "
        "40 Gbps capacity and encrypted VPN tunnels as tertiary backup."
    ))

    doc.add_page_break()

    # ====== PAGE 4: Performance Metrics ======
    add_heading_styled(doc, "3. Performance Metrics", level=1)
    add_body_paragraph(doc, (
        "The following metrics represent system performance across all production environments "
        "during Q1 2025. All values are measured against our SLA targets established in the "
        "2025 Engineering Charter."
    ))

    # Small table that fits portrait
    add_heading_styled(doc, "3.1 Response Time Summary", level=2)
    perf_table = doc.add_table(rows=6, cols=4)
    perf_table.style = "Table Grid"
    perf_headers = ["Service", "P50 (ms)", "P95 (ms)", "P99 (ms)"]
    for i, h in enumerate(perf_headers):
        cell = perf_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    perf_data = [
        ["API Gateway", "1.2", "4.8", "12.3"],
        ["Auth Service", "8.5", "23.1", "67.4"],
        ["Data Pipeline", "45.2", "128.7", "342.1"],
        ["Search Index", "12.8", "34.5", "89.2"],
        ["Notification Hub", "3.1", "9.7", "28.6"],
    ]
    for r, row_data in enumerate(perf_data, 1):
        for c, val in enumerate(row_data):
            perf_table.cell(r, c).text = val

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, (
        "Overall system health remained above 99.95% throughout the quarter. The Data Pipeline "
        "P99 spike in February was attributed to the Aurora cluster migration and has since "
        "been resolved. Current P99 latency is 198ms, well within our 500ms SLA target."
    ))
    add_body_paragraph(doc, (
        "Memory utilization across the fleet averaged 67.3%, with peak usage reaching 84.1% "
        "during the February batch processing window. CPU utilization averaged 42.8% with "
        "headroom for projected Q2 traffic increases of 15-20%."
    ))

    doc.add_page_break()

    # ====== PAGE 5: Wide Infrastructure Capacity Table ======
    add_heading_styled(doc, "4. Infrastructure Capacity Planning", level=1)
    add_body_paragraph(doc, (
        "The table below presents a detailed capacity analysis across all regional data centers, "
        "including current utilization, projected growth, and planned expansions for Q2-Q4 2025."
    ))

    # Create a WIDE table with many columns that will be cut off in portrait
    wide_table = doc.add_table(rows=13, cols=12)
    wide_table.style = "Table Grid"

    wide_headers = [
        "Region", "Data Center", "Rack Units", "CPU Cores",
        "RAM (TB)", "Storage (PB)", "Network (Gbps)",
        "Current Util %", "Projected Q2 %", "Projected Q3 %",
        "Expansion Plan", "Budget ($K)"
    ]
    for i, h in enumerate(wide_headers):
        cell = wide_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    wide_data = [
        ["US-East", "Virginia DC-1", "2,400", "48,000", "192", "4.8", "100",
         "78.3", "85.1", "91.2", "Add 600 RU Q2", "2,450"],
        ["US-East", "Virginia DC-2", "1,800", "36,000", "144", "3.6", "100",
         "65.7", "72.4", "78.9", "GPU cluster Q3", "3,800"],
        ["US-West", "Oregon DC-1", "2,100", "42,000", "168", "4.2", "100",
         "71.2", "77.8", "84.3", "Network upgrade", "1,200"],
        ["US-West", "California DC-1", "1,200", "24,000", "96", "2.4", "40",
         "82.1", "88.5", "94.7", "Migrate to OR", "890"],
        ["EU-West", "Frankfurt DC-1", "1,600", "32,000", "128", "3.2", "100",
         "45.3", "58.7", "71.2", "Scale compute", "1,750"],
        ["EU-West", "Dublin DC-1", "800", "16,000", "64", "1.6", "40",
         "52.8", "64.1", "73.5", "Storage expand", "950"],
        ["APAC", "Singapore DC-1", "1,400", "28,000", "112", "2.8", "100",
         "38.9", "52.3", "65.8", "Full buildout", "4,200"],
        ["APAC", "Tokyo DC-1", "1,000", "20,000", "80", "2.0", "40",
         "61.4", "69.2", "76.8", "Add caching", "680"],
        ["APAC", "Sydney DC-1", "600", "12,000", "48", "1.2", "40",
         "44.7", "55.3", "66.1", "CDN expansion", "520"],
        ["LATAM", "Sao Paulo DC-1", "900", "18,000", "72", "1.8", "40",
         "33.2", "48.6", "62.4", "Initial deploy", "5,100"],
        ["LATAM", "Mexico DC-1", "400", "8,000", "32", "0.8", "20",
         "28.5", "41.2", "53.8", "Network link", "340"],
        ["TOTAL", "", "13,200", "284,000", "1,136", "28.4", "720",
         "57.5", "64.8", "74.4", "", "21,880"],
    ]
    for r, row_data in enumerate(wide_data, 1):
        for c, val in enumerate(row_data):
            cell = wide_table.cell(r, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, (
        "Note: The California DC-1 facility is scheduled for decommissioning by Q4 2025. "
        "All workloads will be migrated to the Oregon DC-1 facility, which offers better "
        "power efficiency (PUE 1.08 vs 1.24) and lower operational costs."
    ))

    doc.add_page_break()

    # ====== PAGE 6: Security Audit ======
    add_heading_styled(doc, "5. Security Audit Results", level=1)
    add_body_paragraph(doc, (
        "The Q1 2025 security audit was conducted by CyberShield Partners from February 10-28. "
        "The assessment covered penetration testing, code review, infrastructure hardening, "
        "and compliance verification against SOC 2 Type II, ISO 27001, and GDPR requirements."
    ))
    add_heading_styled(doc, "5.1 Vulnerability Assessment", level=2)
    vuln_items = [
        "Critical vulnerabilities found: 0 (down from 2 in Q4 2024)",
        "High severity issues: 3 (all remediated within 48 hours)",
        "Medium severity issues: 12 (8 remediated, 4 scheduled for April)",
        "Low severity issues: 27 (tracked in Jira, quarterly remediation)",
        "Informational findings: 45 (documented for awareness)",
    ]
    for item in vuln_items:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_styled(doc, "5.2 Compliance Status", level=2)
    add_body_paragraph(doc, (
        "SOC 2 Type II certification was renewed on March 15, 2025, with zero non-conformities. "
        "ISO 27001 annual surveillance audit is scheduled for May 2025. GDPR compliance review "
        "confirmed full adherence to data processing agreements across all EU operations."
    ))

    doc.add_page_break()

    # ====== PAGE 7: Deployment Timeline ======
    add_heading_styled(doc, "6. Deployment Timeline", level=1)
    add_body_paragraph(doc, (
        "The following timeline outlines major deployment milestones for the remainder of 2025. "
        "All dates are subject to change based on testing outcomes and stakeholder approval."
    ))
    timeline_table = doc.add_table(rows=9, cols=4)
    timeline_table.style = "Table Grid"
    tl_headers = ["Phase", "Target Date", "Description", "Status"]
    for i, h in enumerate(tl_headers):
        cell = timeline_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    tl_data = [
        ["Phase 3A", "Apr 15, 2025", "K8s migration completion", "On Track"],
        ["Phase 3B", "May 1, 2025", "Legacy system decommission", "On Track"],
        ["Phase 4A", "Jun 15, 2025", "Aurora v2 deployment", "Planning"],
        ["Phase 4B", "Jul 30, 2025", "Global CDN rollout", "Planning"],
        ["Phase 5A", "Sep 1, 2025", "ML pipeline integration", "Scoping"],
        ["Phase 5B", "Oct 15, 2025", "Edge computing expansion", "Scoping"],
        ["Phase 6", "Nov 30, 2025", "Full platform v4.0 release", "Roadmap"],
        ["Maintenance", "Dec 15, 2025", "Year-end freeze & audit", "Scheduled"],
    ]
    for r, row_data in enumerate(tl_data, 1):
        for c, val in enumerate(row_data):
            timeline_table.cell(r, c).text = val

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, (
        "The Phase 3A milestone represents the most critical near-term deliverable. The "
        "Kubernetes migration team, led by Senior Engineer Priya Sharma, has completed "
        "migration of 78% of production services with zero downtime incidents."
    ))

    doc.add_page_break()

    # ====== PAGE 8: Budget ======
    add_heading_styled(doc, "7. Budget Allocation", level=1)
    add_body_paragraph(doc, (
        "The 2025 engineering budget of $48.2M has been allocated across the following categories. "
        "Q1 actual spend is $11.4M, which is 2.3% under the projected $11.67M quarterly allocation."
    ))
    budget_table = doc.add_table(rows=8, cols=4)
    budget_table.style = "Table Grid"
    b_headers = ["Category", "Annual Budget ($M)", "Q1 Actual ($M)", "Variance"]
    for i, h in enumerate(b_headers):
        cell = budget_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    b_data = [
        ["Infrastructure & Hardware", "18.5", "4.32", "-3.2%"],
        ["Cloud Services (AWS/GCP)", "12.8", "3.15", "+1.1%"],
        ["Software Licenses", "4.2", "1.05", "0.0%"],
        ["Security & Compliance", "3.6", "0.87", "-2.8%"],
        ["Personnel & Training", "6.1", "1.42", "-5.1%"],
        ["R&D / Innovation", "2.0", "0.45", "-10.0%"],
        ["Contingency Reserve", "1.0", "0.14", "-44.0%"],
    ]
    for r, row_data in enumerate(b_data, 1):
        for c, val in enumerate(row_data):
            budget_table.cell(r, c).text = val

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, (
        "The under-spend in Q1 is primarily driven by delayed hardware procurement for the "
        "Singapore and Sao Paulo data centers. Purchase orders have been approved and deliveries "
        "are expected in April, which will normalize spend by Q2."
    ))

    doc.add_page_break()

    # ====== PAGE 9: Risk Assessment ======
    add_heading_styled(doc, "8. Risk Assessment", level=1)
    add_body_paragraph(doc, (
        "The following risk matrix identifies key technical and operational risks for the "
        "remainder of 2025. Each risk has been evaluated for likelihood and impact, with "
        "mitigation strategies assigned to responsible team leads."
    ))
    risk_table = doc.add_table(rows=7, cols=4)
    risk_table.style = "Table Grid"
    r_headers = ["Risk", "Likelihood", "Impact", "Mitigation"]
    for i, h in enumerate(r_headers):
        cell = risk_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    r_data = [
        ["Supply chain delays", "High", "Medium", "Multi-vendor strategy"],
        ["Key personnel departure", "Medium", "High", "Knowledge documentation"],
        ["Zero-day vulnerability", "Low", "Critical", "Incident response plan"],
        ["Cloud provider outage", "Low", "High", "Multi-cloud failover"],
        ["Regulatory changes (EU)", "Medium", "Medium", "Legal team monitoring"],
        ["Budget overrun", "Low", "Medium", "Monthly variance review"],
    ]
    for r, row_data in enumerate(r_data, 1):
        for c, val in enumerate(row_data):
            risk_table.cell(r, c).text = val

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, (
        "The highest-priority risk remains supply chain volatility for networking equipment. "
        "Lead times for high-capacity switches have increased from 8 weeks to 14 weeks. The "
        "procurement team has established relationships with three alternative suppliers and "
        "maintains a 6-week buffer inventory for critical components."
    ))
    add_body_paragraph(doc, (
        "Personnel risk has been elevated following industry-wide compensation adjustments. "
        "Our retention strategy includes enhanced equity packages, flexible work arrangements, "
        "and a new technical leadership track announced in February 2025."
    ))

    doc.add_page_break()

    # ====== PAGE 10: Recommendations ======
    add_heading_styled(doc, "9. Recommendations & Next Steps", level=1)
    add_body_paragraph(doc, (
        "Based on the findings in this report, the engineering leadership team recommends "
        "the following strategic initiatives for the next two quarters:"
    ))

    recommendations = [
        ("Accelerate Kubernetes Migration",
         "Increase sprint velocity by 20% to complete Phase 3A by April 10, five days ahead "
         "of schedule. Allocate two additional senior engineers from the platform team."),
        ("Expand APAC Presence",
         "Fast-track the Singapore DC-1 buildout to accommodate projected 45% traffic growth "
         "in the Asia-Pacific region. Target operational status by July 2025."),
        ("Implement AI-Driven Monitoring",
         "Deploy the ML-based anomaly detection system currently in beta to all production "
         "environments. Expected to reduce mean time to detection (MTTD) by 60%."),
        ("Enhance Developer Experience",
         "Launch the internal developer portal (Project Compass) by May 2025, consolidating "
         "documentation, API catalogs, and self-service infrastructure provisioning."),
        ("Strengthen Disaster Recovery",
         "Conduct quarterly DR drills across all regions, targeting a recovery time objective "
         "(RTO) of under 15 minutes for Tier 1 services."),
    ]
    for i, (title, desc) in enumerate(recommendations, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. {title}: ")
        run_num.bold = True
        p.add_run(desc)

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, (
        "These recommendations have been reviewed by the executive steering committee and "
        "approved for implementation. Progress will be tracked in the weekly engineering "
        "standup and reported in the Q2 2025 Engineering Report."
    ))
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, (
        "Respectfully submitted,"
    ))
    add_body_paragraph(doc, (
        "Dr. Elena Vasquez\n"
        "Lead Systems Architect\n"
        "Meridian Systems Engineering"
    ))

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
