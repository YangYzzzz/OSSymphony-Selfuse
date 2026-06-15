"""
Initial Setup: User manual document with headings but no Table of Contents
Task ID: writer_tech_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

# Install python-docx on VM if not present
subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_text(doc, text, space_after=Pt(6)):
    """Add a normal body paragraph."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = space_after
    return para


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Configure heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'

    # ============================================================
    # Chapter 1: Introduction
    # ============================================================
    doc.add_heading('Introduction', level=1)

    add_body_text(doc, (
        'Welcome to the DataStream Analytics Platform User Manual. This comprehensive guide '
        'provides detailed instructions for installing, configuring, and operating the '
        'DataStream Analytics Platform version 4.2. Whether you are a first-time user or an '
        'experienced administrator, this manual covers all the features and capabilities of '
        'the platform.'
    ))

    doc.add_heading('Purpose of This Manual', level=2)
    add_body_text(doc, (
        'This manual is intended for system administrators, data analysts, and business '
        'intelligence professionals who use the DataStream Analytics Platform on a daily basis. '
        'It covers everything from initial setup to advanced configuration scenarios, including '
        'troubleshooting procedures and best practices for optimal performance.'
    ))

    doc.add_heading('Document Conventions', level=3)
    add_body_text(doc, (
        'Throughout this manual, the following conventions are used: Bold text indicates '
        'menu items and button labels. Italic text indicates file names and paths. '
        'Monospaced text indicates command-line input and code samples. Warning boxes '
        'highlight potential issues that may cause data loss or system instability.'
    ))

    doc.add_heading('Revision History', level=3)
    add_body_text(doc, (
        'Version 4.2 was released on March 15, 2025 and includes new dashboard widgets, '
        'improved data pipeline performance, and enhanced security features. Version 4.1 '
        'was released on November 8, 2024 with bug fixes and minor UI improvements. '
        'Version 4.0 was the initial major release on June 1, 2024.'
    ))

    doc.add_heading('System Requirements', level=2)
    add_body_text(doc, (
        'Before installing the DataStream Analytics Platform, ensure your system meets the '
        'minimum hardware and software requirements listed below. Failure to meet these '
        'requirements may result in degraded performance or installation failures.'
    ))

    doc.add_heading('Hardware Requirements', level=3)
    add_body_text(doc, (
        'The minimum hardware requirements are: 8-core CPU (Intel Xeon E5 or equivalent), '
        '32 GB RAM, 500 GB SSD storage for the application and temporary files, and a '
        'dedicated 1 Gbps network interface. For production environments handling more than '
        '10,000 concurrent sessions, we recommend 16-core CPU, 64 GB RAM, and 2 TB NVMe storage.'
    ))

    doc.add_heading('Software Requirements', level=3)
    add_body_text(doc, (
        'Supported operating systems include Ubuntu 22.04 LTS, Red Hat Enterprise Linux 8.x '
        'or 9.x, and Windows Server 2022. The platform requires Java Runtime Environment '
        'version 17 or later, PostgreSQL 15.x, and Redis 7.x for caching. Docker 24.x is '
        'required for containerized deployments.'
    ))

    # ============================================================
    # Chapter 2: Installation
    # ============================================================
    doc.add_heading('Installation', level=1)

    add_body_text(doc, (
        'This chapter walks you through the complete installation process for the DataStream '
        'Analytics Platform. The installation can be performed using either the graphical '
        'installer or the command-line interface. Both methods produce identical results.'
    ))

    doc.add_heading('Pre-Installation Checklist', level=2)
    add_body_text(doc, (
        'Before beginning the installation, verify the following: all hardware meets the '
        'minimum requirements specified in Chapter 1, the target server has network access to '
        'the DataStream license server at license.datastream-analytics.com on port 8443, '
        'PostgreSQL is installed and a blank database has been created, and the installation '
        'user has root or administrator privileges.'
    ))

    doc.add_heading('Downloading the Installer', level=3)
    add_body_text(doc, (
        'Navigate to the DataStream Customer Portal at portal.datastream-analytics.com and '
        'log in with your enterprise credentials. Select the Downloads section, choose your '
        'operating system, and download the appropriate installer package. The SHA-256 checksum '
        'is provided alongside each download for verification purposes.'
    ))

    doc.add_heading('Verifying the Download', level=3)
    add_body_text(doc, (
        'After downloading, verify the integrity of the installer by comparing its SHA-256 '
        'hash against the published checksum. On Linux, use the sha256sum command. On Windows, '
        'use the Get-FileHash PowerShell cmdlet. If the checksums do not match, re-download '
        'the installer and try again.'
    ))

    doc.add_heading('Graphical Installation', level=2)
    add_body_text(doc, (
        'Launch the graphical installer by double-clicking the downloaded package on Windows '
        'or running the installer script with the --gui flag on Linux. The installation wizard '
        'will guide you through six configuration screens: License Agreement, Installation '
        'Directory, Database Connection, Cache Configuration, Network Settings, and Summary.'
    ))

    doc.add_heading('License Agreement Screen', level=3)
    add_body_text(doc, (
        'Read the End User License Agreement carefully. You must accept the terms to proceed '
        'with the installation. The agreement covers usage rights, data handling obligations, '
        'and support entitlements. A copy of the agreement is saved to the installation '
        'directory for future reference.'
    ))

    doc.add_heading('Database Configuration Screen', level=3)
    add_body_text(doc, (
        'Enter the PostgreSQL connection details: hostname or IP address, port number (default '
        '5432), database name, and authentication credentials. The installer will test the '
        'connection and verify that the database is empty before proceeding. If the test fails, '
        'check your firewall rules and PostgreSQL pg_hba.conf settings.'
    ))

    doc.add_heading('Command-Line Installation', level=2)
    add_body_text(doc, (
        'For automated or headless installations, use the command-line installer with a '
        'configuration file. Create a YAML configuration file specifying all installation '
        'parameters, then run the installer with the --config flag. This method is recommended '
        'for deploying to multiple servers using configuration management tools like Ansible '
        'or Puppet.'
    ))

    doc.add_heading('Configuration File Format', level=3)
    add_body_text(doc, (
        'The configuration file uses YAML syntax with the following top-level sections: '
        'license (containing the key and acceptance flag), paths (installation directory and '
        'data directory), database (connection parameters), cache (Redis connection), and '
        'network (bind address and port). Example configuration files are provided in the '
        'installer package under the examples directory.'
    ))

    # ============================================================
    # Chapter 3: Configuration
    # ============================================================
    doc.add_heading('Configuration', level=1)

    add_body_text(doc, (
        'After installation, the DataStream Analytics Platform requires configuration to '
        'match your organization requirements. This chapter covers the main configuration '
        'areas including data sources, user management, security policies, and dashboard '
        'customization.'
    ))

    doc.add_heading('Data Source Management', level=2)
    add_body_text(doc, (
        'The platform supports connections to over 40 data source types including relational '
        'databases (MySQL, PostgreSQL, Oracle, SQL Server), NoSQL databases (MongoDB, '
        'Cassandra, Elasticsearch), cloud storage (AWS S3, Azure Blob, Google Cloud Storage), '
        'and streaming platforms (Apache Kafka, Amazon Kinesis, Google Pub/Sub).'
    ))

    doc.add_heading('Adding a Database Connection', level=3)
    add_body_text(doc, (
        'To add a new database data source, navigate to Administration > Data Sources > Add '
        'New. Select the database type from the dropdown menu, then enter the connection '
        'parameters: hostname, port, database name, schema, and credentials. Use the Test '
        'Connection button to verify connectivity before saving.'
    ))

    doc.add_heading('Configuring Data Refresh Schedules', level=3)
    add_body_text(doc, (
        'Each data source can have an independent refresh schedule. Navigate to the data '
        'source settings and select the Schedule tab. Choose from predefined intervals '
        '(every 15 minutes, hourly, daily) or define a custom cron expression. For real-time '
        'data sources, enable the streaming mode option instead.'
    ))

    doc.add_heading('User and Role Management', level=2)
    add_body_text(doc, (
        'The platform uses role-based access control (RBAC) with four default roles: Viewer, '
        'Analyst, Editor, and Administrator. Each role has a predefined set of permissions '
        'that control access to dashboards, data sources, and administrative functions. Custom '
        'roles can be created to meet specific organizational requirements.'
    ))

    doc.add_heading('Creating User Accounts', level=3)
    add_body_text(doc, (
        'Navigate to Administration > Users > Add User. Enter the user email address, display '
        'name, and assign one or more roles. If LDAP integration is configured, users can be '
        'imported from the directory. New users receive a welcome email with instructions to '
        'set their password and configure two-factor authentication.'
    ))

    doc.add_heading('Permission Matrix', level=3)
    add_body_text(doc, (
        'The permission matrix defines granular access controls for each role. Permissions '
        'are organized into categories: Dashboard (view, create, edit, delete, share), '
        'Data Source (view, create, configure, delete), Report (view, create, schedule, '
        'export), and Administration (user management, system settings, audit logs). '
        'Custom roles can have any combination of these permissions.'
    ))

    doc.add_heading('Security Configuration', level=2)
    add_body_text(doc, (
        'Security is a critical aspect of the DataStream Analytics Platform. This section '
        'covers authentication methods, encryption settings, and audit logging. All security '
        'configurations should be reviewed and approved by your organization IT security team '
        'before deployment to production.'
    ))

    doc.add_heading('Authentication Methods', level=3)
    add_body_text(doc, (
        'The platform supports multiple authentication methods: local database authentication '
        'with password complexity requirements, LDAP/Active Directory integration for '
        'enterprise single sign-on, SAML 2.0 for federated identity providers, and OAuth 2.0 '
        'for third-party application access. Two-factor authentication can be enforced for '
        'all methods using TOTP or hardware security keys.'
    ))

    doc.add_heading('Encryption Settings', level=3)
    add_body_text(doc, (
        'All data in transit is encrypted using TLS 1.3. Data at rest is encrypted using '
        'AES-256-GCM. The encryption keys are stored in a dedicated key management service '
        'and rotated every 90 days automatically. Certificate management supports both '
        'self-signed certificates for development environments and CA-signed certificates '
        'for production deployments.'
    ))

    # ============================================================
    # Chapter 4: Dashboard Operations
    # ============================================================
    doc.add_heading('Dashboard Operations', level=1)

    add_body_text(doc, (
        'Dashboards are the primary interface for visualizing and interacting with your data. '
        'This chapter covers creating, editing, sharing, and managing dashboards in the '
        'DataStream Analytics Platform.'
    ))

    doc.add_heading('Creating a New Dashboard', level=2)
    add_body_text(doc, (
        'To create a new dashboard, click the Create Dashboard button on the home screen or '
        'navigate to Dashboards > New. Enter a title, optional description, and select the '
        'default data source. The dashboard editor opens with a blank canvas where you can '
        'add widgets, filters, and layout sections.'
    ))

    doc.add_heading('Widget Types', level=3)
    add_body_text(doc, (
        'The platform offers 15 widget types organized into four categories. Chart widgets '
        'include line, bar, area, pie, scatter, and radar charts. Table widgets include '
        'standard data tables and pivot tables. Metric widgets include single value indicators, '
        'gauges, and sparklines. Map widgets include choropleth maps and point maps. Each '
        'widget type has specific configuration options detailed in the Widget Reference appendix.'
    ))

    doc.add_heading('Adding and Configuring Widgets', level=3)
    add_body_text(doc, (
        'Drag a widget type from the widget palette onto the dashboard canvas. The widget '
        'configuration panel opens automatically. Select the data source, define the query '
        'or select fields, and configure visualization options such as colors, labels, and '
        'legends. Preview the widget in real-time as you make changes. Click Apply to save '
        'the widget configuration.'
    ))

    doc.add_heading('Dashboard Sharing and Permissions', level=2)
    add_body_text(doc, (
        'Dashboards can be shared with individual users, groups, or made public within the '
        'organization. Navigate to the dashboard settings and select the Sharing tab. Choose '
        'the access level for each recipient: Viewer (read-only), Editor (can modify widgets '
        'and layout), or Owner (full control including deletion and permission management).'
    ))

    doc.add_heading('Export and Scheduling', level=3)
    add_body_text(doc, (
        'Dashboards can be exported to PDF, PNG, or interactive HTML format. To schedule '
        'automatic exports, navigate to the dashboard settings and select the Schedule tab. '
        'Configure the export format, frequency, and email recipients. Scheduled exports are '
        'generated during off-peak hours to minimize system impact.'
    ))

    doc.add_heading('Dashboard Templates', level=3)
    add_body_text(doc, (
        'Save frequently used dashboard layouts as templates for reuse. Navigate to Dashboard '
        'Settings > Save as Template. Templates preserve the layout, widget types, and '
        'configuration but not the bound data sources. When creating a new dashboard from a '
        'template, you will be prompted to map each template data source to an available '
        'connection in your environment.'
    ))

    # ============================================================
    # Chapter 5: Data Pipeline Management
    # ============================================================
    doc.add_heading('Data Pipeline Management', level=1)

    add_body_text(doc, (
        'Data pipelines automate the extraction, transformation, and loading (ETL) of data '
        'from source systems into the analytics platform. This chapter covers pipeline '
        'creation, monitoring, error handling, and optimization techniques.'
    ))

    doc.add_heading('Pipeline Architecture', level=2)
    add_body_text(doc, (
        'Each pipeline consists of three stages: extraction (pulling data from source systems), '
        'transformation (cleaning, enriching, and reshaping data), and loading (writing '
        'processed data to the analytics data store). Pipelines execute as directed acyclic '
        'graphs (DAGs) allowing parallel processing of independent stages.'
    ))

    doc.add_heading('Extraction Stage', level=3)
    add_body_text(doc, (
        'The extraction stage supports full and incremental extraction modes. Full extraction '
        'retrieves the complete dataset on each run, suitable for small reference tables. '
        'Incremental extraction tracks changes using timestamps, sequence numbers, or change '
        'data capture (CDC) logs, significantly reducing data transfer volumes for large tables.'
    ))

    doc.add_heading('Transformation Stage', level=3)
    add_body_text(doc, (
        'Transformations are defined using a visual node editor or SQL expressions. Built-in '
        'transformation nodes include: Filter (row-level conditions), Aggregate (grouping and '
        'summarization), Join (combining data from multiple sources), Pivot (reshaping data), '
        'Lookup (enrichment from reference tables), and Custom (Python or SQL scripts). '
        'Transformations execute in parallel where dependencies allow.'
    ))

    doc.add_heading('Pipeline Monitoring', level=2)
    add_body_text(doc, (
        'The pipeline monitoring dashboard provides real-time visibility into pipeline '
        'execution status, data volumes, processing times, and error rates. Each pipeline '
        'run is logged with detailed metrics including rows processed, bytes transferred, '
        'stage durations, and any warnings or errors encountered.'
    ))

    doc.add_heading('Setting Up Alerts', level=3)
    add_body_text(doc, (
        'Configure alerts to notify the operations team when pipeline issues occur. Navigate '
        'to Pipelines > Alerts > Create Alert. Define the trigger condition (e.g., pipeline '
        'failure, processing time exceeds threshold, data volume anomaly), notification '
        'channels (email, Slack, PagerDuty), and escalation rules. Alert history is maintained '
        'for 90 days for audit purposes.'
    ))

    doc.add_heading('Error Handling and Recovery', level=3)
    add_body_text(doc, (
        'When a pipeline stage fails, the platform automatically retries the failed stage up '
        'to three times with exponential backoff. If all retries fail, the pipeline is marked '
        'as failed and an alert is triggered. Failed pipelines can be manually restarted from '
        'the failed stage or from the beginning. The platform maintains checkpoint data to '
        'support efficient recovery without reprocessing successfully completed stages.'
    ))

    # ============================================================
    # Chapter 6: Troubleshooting
    # ============================================================
    doc.add_heading('Troubleshooting', level=1)

    add_body_text(doc, (
        'This chapter provides guidance for diagnosing and resolving common issues with the '
        'DataStream Analytics Platform. For issues not covered in this chapter, contact '
        'DataStream Technical Support at support@datastream-analytics.com or call the 24/7 '
        'support hotline at +1-888-555-0199.'
    ))

    doc.add_heading('Common Installation Issues', level=2)
    add_body_text(doc, (
        'The most common installation issues are related to database connectivity, insufficient '
        'disk space, and permission errors. This section provides step-by-step resolution '
        'procedures for each issue type.'
    ))

    doc.add_heading('Database Connection Failures', level=3)
    add_body_text(doc, (
        'If the installer reports a database connection failure, verify the following: the '
        'PostgreSQL service is running (systemctl status postgresql), the database host is '
        'reachable from the application server (telnet dbhost 5432), the database user has '
        'the required privileges (CREATEDB, CONNECT), and the pg_hba.conf file allows '
        'connections from the application server IP address. After making changes to '
        'pg_hba.conf, reload the PostgreSQL configuration.'
    ))

    doc.add_heading('Insufficient Disk Space', level=3)
    add_body_text(doc, (
        'The installer requires at least 500 MB of free space in the installation directory '
        'and 2 GB in the temporary directory. Check available space using df -h on Linux or '
        'the Disk Management tool on Windows. If space is insufficient, either free up space '
        'by removing unnecessary files or specify an alternative installation directory with '
        'adequate capacity.'
    ))

    doc.add_heading('Performance Issues', level=2)
    add_body_text(doc, (
        'Performance degradation can occur due to insufficient resources, misconfigured '
        'settings, or excessive data volumes. This section covers the most common performance '
        'issues and their resolutions.'
    ))

    doc.add_heading('Slow Dashboard Loading', level=3)
    add_body_text(doc, (
        'If dashboards take more than 10 seconds to load, check the following: verify that '
        'the Redis cache is operational and has sufficient memory allocated, review widget '
        'queries for inefficient joins or missing indexes, check the browser developer tools '
        'network tab for slow API responses, and ensure the application server has adequate '
        'CPU and memory resources. Consider enabling query caching for frequently accessed '
        'dashboards.'
    ))

    doc.add_heading('High Memory Usage', level=3)
    add_body_text(doc, (
        'The platform Java process may consume excessive memory under heavy load. Review the '
        'JVM heap settings in the configuration file (default: -Xmx4g). For servers with '
        'more than 32 GB RAM, increase the heap to 8 GB or 16 GB. Monitor garbage collection '
        'frequency and duration using the built-in JMX metrics. If GC pauses exceed 500 ms, '
        'consider switching to the ZGC garbage collector.'
    ))

    doc.add_heading('Log Files and Diagnostics', level=2)
    add_body_text(doc, (
        'The platform writes diagnostic information to several log files in the logs directory. '
        'The application.log file contains general operational messages. The error.log file '
        'records exceptions and stack traces. The access.log file records all HTTP requests. '
        'The pipeline.log file records data pipeline execution details. Log rotation is '
        'configured to retain 30 days of log history with daily rotation and gzip compression.'
    ))

    doc.add_heading('Collecting Support Bundles', level=3)
    add_body_text(doc, (
        'When contacting Technical Support, generate a support bundle using the Administration '
        '> Diagnostics > Generate Support Bundle function. The bundle includes sanitized '
        'configuration files, recent log excerpts, system resource metrics, and database '
        'schema information. Sensitive data such as passwords and API keys are automatically '
        'redacted. Upload the bundle to the support portal or attach it to your support ticket.'
    ))

    doc.add_heading('Interpreting Log Messages', level=3)
    add_body_text(doc, (
        'Log messages follow the format: [TIMESTAMP] [LEVEL] [COMPONENT] Message. Log levels '
        'are DEBUG, INFO, WARN, ERROR, and FATAL. Component identifiers include WEB (HTTP '
        'layer), PIPE (pipeline engine), CACHE (Redis integration), DB (database layer), and '
        'AUTH (authentication). Search for ERROR and FATAL messages first when diagnosing '
        'issues, then examine surrounding INFO and WARN messages for context.'
    ))

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
