"""
Initial Setup: Insert section named 'Appendix A' around last three paragraphs
Task ID: writer_struct_011
Domain: libreoffice_writer

Creates annual_report.docx at ~/Desktop/ with a 4-page report.
No sections exist in the initial file.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'annual_report'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Cover / Title ---
    title = doc.add_heading('Annual Report 2024', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('Meridian Technologies Corporation')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(14)
    run.bold = True

    fiscal = doc.add_paragraph('Fiscal Year Ending December 31, 2024')
    fiscal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fiscal.runs[0].font.size = Pt(12)

    doc.add_paragraph('')

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'Meridian Technologies Corporation delivered strong performance in fiscal year 2024, '
        'achieving record revenues of $284.7 million, a 12.3% increase over the prior year. '
        'Net income reached $41.2 million, representing a net margin of 14.5%. '
        'These results reflect our continued investment in product innovation, '
        'customer success initiatives, and operational excellence across all business units.'
    )
    doc.add_paragraph(
        'Our cloud services division led growth with a 28% year-over-year increase, '
        'while our legacy hardware segment demonstrated resilience with stable margins '
        'despite market headwinds. The Board of Directors approved a dividend of $0.45 per share, '
        'reflecting confidence in the company\'s financial health and future prospects.'
    )

    # --- Section 2: Financial Highlights ---
    doc.add_heading('2. Financial Highlights', level=1)
    doc.add_paragraph(
        'Total revenue increased from $253.5 million in 2023 to $284.7 million in 2024. '
        'Operating expenses were $213.8 million, yielding operating income of $70.9 million. '
        'EBITDA for the year was $89.4 million, compared to $76.2 million in 2023. '
        'Cash and equivalents stood at $62.1 million at year end, with total assets of $387.5 million.'
    )

    # Financial table
    doc.add_paragraph('Key Financial Metrics ($ millions)').runs[0].bold = True
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Metric', '2024', '2023']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    rows_data = [
        ['Total Revenue', '284.7', '253.5'],
        ['Operating Income', '70.9', '61.4'],
        ['Net Income', '41.2', '35.8'],
        ['EBITDA', '89.4', '76.2'],
        ['EPS (diluted)', '$1.82', '$1.59'],
    ]
    for row_idx, row_data in enumerate(rows_data, 1):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx, col_idx).text = val

    doc.add_paragraph('')

    # --- Section 3: Operations Review ---
    doc.add_heading('3. Operations Review', level=1)
    doc.add_paragraph(
        'The cloud services division achieved $121.4 million in revenue, representing 42.6% '
        'of total company revenue. Key contributors included the launch of MeridianCloud Pro '
        'in Q2 2024, which signed 847 enterprise customers by year end, and the expansion of '
        'our data center footprint in Singapore and Frankfurt.'
    )
    doc.add_paragraph(
        'Hardware solutions revenue was $98.3 million, with gross margins improving 1.8 percentage '
        'points to 34.2% through supply chain optimization. Professional services contributed '
        '$65.0 million, driven by strong demand for integration and migration support.'
    )

    doc.add_page_break()

    # --- Section 4: Strategy and Outlook ---
    doc.add_heading('4. Strategy and Outlook', level=1)
    doc.add_paragraph(
        'Looking ahead to 2025, Meridian Technologies is focused on three strategic pillars: '
        'accelerating cloud growth, deepening customer relationships, and expanding into '
        'adjacent markets. We expect total revenue in the range of $315 million to $330 million, '
        'representing growth of 11% to 16% over 2024.'
    )
    doc.add_paragraph(
        'Investment in research and development will increase to $38.5 million, up 22% from 2024, '
        'as we accelerate development of our next-generation AI-assisted analytics platform. '
        'We also plan to expand headcount by approximately 280 employees globally, '
        'with a focus on engineering, sales, and customer success roles.'
    )

    # --- Section 5: Corporate Governance ---
    doc.add_heading('5. Corporate Governance', level=1)
    doc.add_paragraph(
        'The Board of Directors consists of eight members, six of whom are independent. '
        'During 2024, the Audit Committee oversaw a comprehensive review of internal controls, '
        'resulting in several process improvements. The Compensation Committee benchmarked '
        'executive pay against 18 peer companies and made adjustments to align with market practice.'
    )
    doc.add_paragraph(
        'Environmental, Social, and Governance (ESG) commitments remain a priority. '
        'Carbon emissions were reduced by 9% versus 2023 through data center efficiency programs. '
        'Employee engagement scores reached an all-time high of 78%, and diversity hiring '
        'increased female representation in technical roles to 34%.'
    )

    doc.add_page_break()

    # --- Section 6: Risk Factors ---
    doc.add_heading('6. Risk Factors', level=1)
    doc.add_paragraph(
        'The company faces several material risks including competitive pressure from larger cloud '
        'providers, foreign exchange volatility given that 31% of revenue is denominated in '
        'non-USD currencies, and potential disruption from evolving data privacy regulations '
        'in the European Union and Asia-Pacific markets.'
    )
    doc.add_paragraph(
        'Cybersecurity risks are actively managed through our ISO 27001-certified security program, '
        'which includes regular penetration testing, a 24/7 security operations center, and '
        'a cyber liability insurance policy with $50 million in coverage.'
    )

    # --- Section 7: Auditor Report ---
    doc.add_heading('7. Independent Auditor\'s Report', level=1)
    doc.add_paragraph(
        'To the Shareholders and Board of Directors of Meridian Technologies Corporation: '
        'We have audited the accompanying consolidated financial statements of Meridian '
        'Technologies Corporation as of December 31, 2024 and 2023. In our opinion, the '
        'financial statements present fairly, in all material respects, the financial position '
        'of the Company in conformity with U.S. Generally Accepted Accounting Principles.'
    )

    doc.add_page_break()

    # --- Appendix heading paragraph (NOT part of the appendix content section) ---
    doc.add_heading('Appendix', level=1)

    # --- Last three paragraphs: appendix content (no section wrapper yet) ---
    doc.add_paragraph(
        'Table A-1 shows the raw revenue data by product line for fiscal years 2022, 2023, and 2024. '
        'Cloud Services generated $94.8M in 2022, $94.8M in 2023 growing to $121.4M in 2024. '
        'Hardware Solutions contributed $107.2M, $104.9M, and $98.3M in respective years. '
        'Professional Services grew from $51.3M to $54.3M to $65.0M over the same period.'
    )
    doc.add_paragraph(
        'Table A-2 presents quarterly revenue breakdown for fiscal year 2024. Q1 revenue was '
        '$65.3M, Q2 revenue was $69.8M, Q3 revenue was $73.2M, and Q4 revenue was $76.4M. '
        'The sequential growth pattern reflects seasonal demand trends and the mid-year launch '
        'of MeridianCloud Pro which significantly boosted Q3 and Q4 performance.'
    )
    doc.add_paragraph(
        'The financial data presented in Tables A-1 and A-2 has been derived from audited '
        'consolidated financial statements and has not been adjusted for any subsequent events. '
        'All figures are in millions of U.S. dollars unless otherwise noted. '
        'Certain prior year amounts have been reclassified to conform to the current year presentation.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
