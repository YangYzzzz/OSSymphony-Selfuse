"""
Initial Setup: Create a Writer document with a glossary list of 30 terms in single-column paragraph format.
Task ID: writer_fs_089
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_089'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Technical Reference Manual', level=0)

    # Introductory paragraph
    intro = doc.add_paragraph(
        'This document provides a comprehensive glossary of terms '
        'commonly used in software engineering and computer science. '
        'The glossary below lists key terminology that practitioners '
        'should be familiar with.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # 30 glossary terms in alphabetical order, presented as simple paragraphs
    # Each term is bold followed by its definition
    glossary_terms = [
        ("API (Application Programming Interface)",
         "A set of protocols and tools that allows different software applications to communicate with each other."),
        ("Agile",
         "An iterative approach to software development that emphasizes flexibility, collaboration, and customer feedback."),
        ("Binary Tree",
         "A hierarchical data structure in which each node has at most two children, referred to as left and right."),
        ("Buffer Overflow",
         "A condition where a program writes data beyond the allocated memory buffer, potentially causing crashes or security vulnerabilities."),
        ("Cache",
         "A high-speed storage layer that stores a subset of data so future requests can be served faster."),
        ("Concurrency",
         "The ability of a system to handle multiple tasks simultaneously by interleaving their execution."),
        ("Deadlock",
         "A situation where two or more processes are unable to proceed because each is waiting for the other to release resources."),
        ("Dependency Injection",
         "A design pattern where objects receive their dependencies from external sources rather than creating them internally."),
        ("Encryption",
         "The process of converting plaintext into ciphertext to protect data from unauthorized access."),
        ("Event-Driven Architecture",
         "A software design pattern where the flow of the program is determined by events such as user actions or sensor outputs."),
        ("Firewall",
         "A network security system that monitors and controls incoming and outgoing traffic based on predefined rules."),
        ("Functional Programming",
         "A programming paradigm that treats computation as the evaluation of mathematical functions and avoids mutable state."),
        ("Garbage Collection",
         "An automatic memory management process that reclaims memory occupied by objects no longer in use."),
        ("Git",
         "A distributed version control system used for tracking changes in source code during software development."),
        ("Hash Table",
         "A data structure that maps keys to values using a hash function for efficient lookup operations."),
        ("Hypervisor",
         "Software that creates and manages virtual machines, allowing multiple operating systems to share hardware resources."),
        ("Idempotent",
         "An operation that produces the same result regardless of how many times it is executed with the same input."),
        ("Inheritance",
         "An object-oriented programming mechanism where a class derives properties and behaviors from a parent class."),
        ("JSON (JavaScript Object Notation)",
         "A lightweight data interchange format that is easy for humans to read and write and for machines to parse."),
        ("JIT Compilation",
         "A technique where bytecode is compiled into native machine code at runtime to improve execution performance."),
        ("Kubernetes",
         "An open-source container orchestration platform for automating deployment, scaling, and management of containerized applications."),
        ("Latency",
         "The time delay between a request being made and the corresponding response being received."),
        ("Microservices",
         "An architectural style that structures an application as a collection of loosely coupled, independently deployable services."),
        ("Mutex",
         "A synchronization primitive used to prevent multiple threads from simultaneously accessing a shared resource."),
        ("Normalization",
         "The process of organizing database tables to reduce redundancy and improve data integrity."),
        ("OAuth",
         "An open standard for access delegation commonly used to grant websites limited access to user information without exposing passwords."),
        ("Polymorphism",
         "The ability of different objects to respond to the same method call in different ways based on their type."),
        ("Race Condition",
         "A flaw where the output of a process depends on the timing or sequence of uncontrollable events."),
        ("Serialization",
         "The process of converting an object or data structure into a format suitable for storage or transmission."),
        ("Thread Pool",
         "A collection of pre-instantiated reusable threads that can be used to execute tasks without the overhead of creating new threads."),
    ]

    for term, definition in glossary_terms:
        para = doc.add_paragraph()
        run_term = para.add_run(term)
        run_term.bold = True
        run_def = para.add_run(f" - {definition}")
        para.paragraph_format.space_after = Pt(4)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
