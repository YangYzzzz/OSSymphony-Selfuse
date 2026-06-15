"""
Reward Script: Create a 'Glossary' section with 3-column layout and alphabetical separators
Task ID: writer_fs_089
Domain: libreoffice_writer
Scoring:
  Component 1: 'Glossary' heading exists as Heading 1 style (0.15)
  Component 2: 3-column layout section exists (0.25)
  Component 3: Column spacing approximately 0.4 cm (144000 EMU) (0.15)
  Component 4: Alphabetical bold letter separators present (0.30)
  Component 5: Terms correctly grouped under letter headings (0.15)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_089'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 'Glossary' heading exists with Heading 1 style (0.15 points)
    # Initial file has NO Heading 1 paragraphs, golden has 'Glossary' as Heading 1
    try:
        glossary_headings = [
            p for p in doc.paragraphs
            if p.style.name == 'Heading 1' and 'glossary' in p.text.strip().lower()
        ]
        if len(glossary_headings) > 0:
            print(f"PASS: Component 1 — 'Glossary' heading found with Heading 1 style (0.15 pts)")
            total_score += 0.15
        else:
            # Also check Heading 2, Heading 3 as partial
            glossary_any_heading = [
                p for p in doc.paragraphs
                if p.style.name.startswith('Heading') and 'glossary' in p.text.strip().lower()
            ]
            if glossary_any_heading:
                print(f"FAIL: Component 1 — 'Glossary' found as {glossary_any_heading[0].style.name}, expected Heading 1")
            else:
                print("FAIL: Component 1 — No 'Glossary' heading found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 3-column layout section exists (0.25 points)
    # Initial file has only 1 section with default single column
    # Golden file has a section with w:cols num=3
    try:
        col3_section_idx = -1
        for i, section in enumerate(doc.sections):
            sectPr = section._sectPr
            cols_elem = sectPr.find(qn('w:cols'))
            if cols_elem is not None:
                num_cols = cols_elem.get(qn('w:num'))
                if num_cols is not None and int(num_cols) == 3:
                    col3_section_idx = i
                    break
        if col3_section_idx >= 0:
            print(f"PASS: Component 2 — Found 3-column section (section {col3_section_idx}) (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 2 — No 3-column section found in any section")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Column spacing approximately 0.4 cm (144000 EMU) (0.15 points)
    # Initial file has no multi-column section, so this check only passes on golden
    try:
        spacing_score = 0.0
        for section in doc.sections:
            sectPr = section._sectPr
            cols_elem = sectPr.find(qn('w:cols'))
            if cols_elem is not None:
                num_cols = cols_elem.get(qn('w:num'))
                if num_cols is not None and int(num_cols) >= 3:
                    space_val = cols_elem.get(qn('w:space'))
                    if space_val is not None:
                        space_int = int(space_val)
                        # 0.4 cm = 144000 EMU; allow some tolerance (100000-200000)
                        if 100000 <= space_int <= 200000:
                            spacing_score = 0.15
                            print(f"PASS: Component 3 — Column spacing {space_int} EMU (~{space_int/360000:.2f} cm) (0.15 pts)")
                            break
                    # Check individual column spacing if equalWidth is not set
                    if spacing_score == 0.0:
                        col_elems = cols_elem.findall(qn('w:col'))
                        for col in col_elems:
                            col_space = col.get(qn('w:space'))
                            if col_space and 100000 <= int(col_space) <= 200000:
                                spacing_score = 0.15
                                print(f"PASS: Component 3 — Individual column spacing {col_space} EMU (0.15 pts)")
                                break
                        if spacing_score > 0:
                            break
        if spacing_score > 0:
            total_score += spacing_score
        else:
            print("FAIL: Component 3 — Column spacing not approximately 0.4 cm in any 3-column section")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Alphabetical bold letter separators present (0.30 points)
    # Initial file has ZERO single-letter bold paragraphs
    # Golden file has 19 bold letter separators (A through T, skipping Q)
    try:
        alpha_separators = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if len(text) == 1 and text.isalpha() and text.isupper():
                # Check if the letter run is bold
                is_bold = any(r.bold for r in p.runs if r.text.strip())
                alpha_separators.append((text, is_bold))

        unique_letters = set(l for l, b in alpha_separators)
        bold_letters = set(l for l, b in alpha_separators if b)

        # Need at least 10 unique letter separators to get any credit
        if len(unique_letters) >= 10 and len(bold_letters) >= len(unique_letters) * 0.8:
            # Full credit: enough separators and most are bold
            print(f"PASS: Component 4 — Found {len(unique_letters)} letter separators, {len(bold_letters)} bold (0.30 pts)")
            total_score += 0.30
        elif len(unique_letters) >= 10:
            # Separators present but not all bold - partial credit
            earned = round(0.15 + 0.15 * (len(bold_letters) / max(len(unique_letters), 1)), 2) if len(bold_letters) > 0 else 0.15
            print(f"PARTIAL: Component 4 — Found {len(unique_letters)} separators, {len(bold_letters)} bold ({earned} pts)")
            total_score += earned
        elif len(unique_letters) >= 5:
            earned = round(0.10 * (len(unique_letters) / 19), 2)
            print(f"PARTIAL: Component 4 — Found {len(unique_letters)} letter separators ({earned} pts)")
            total_score += earned
        else:
            print(f"FAIL: Component 4 — Found only {len(unique_letters)} letter separators (need >= 10)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Terms correctly grouped under letter headings (0.15 points)
    # Verify that terms appear after the correct letter heading
    # Initial file has no letter headings, so this only passes on golden
    try:
        # Build a mapping: current_letter -> terms that follow
        current_letter = None
        letter_terms = {}
        correct_groupings = 0
        total_terms = 0

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            # Check if it's a letter separator
            if len(text) == 1 and text.isalpha() and text.isupper():
                current_letter = text
                if current_letter not in letter_terms:
                    letter_terms[current_letter] = []
                continue
            # Check if it's a glossary term (contains ' - ')
            if current_letter and ' - ' in text:
                # The term should start with the current letter
                term_first_char = text[0].upper()
                total_terms += 1
                if term_first_char == current_letter:
                    correct_groupings += 1
                    letter_terms[current_letter].append(text[:30])

        if total_terms > 0 and correct_groupings >= total_terms * 0.8:
            print(f"PASS: Component 5 — {correct_groupings}/{total_terms} terms correctly grouped ({0.15} pts)")
            total_score += 0.15
        elif total_terms > 0 and correct_groupings > 0:
            earned = round(0.15 * (correct_groupings / total_terms), 2)
            print(f"PARTIAL: Component 5 — {correct_groupings}/{total_terms} terms correctly grouped ({earned} pts)")
            total_score += earned
        else:
            print(f"FAIL: Component 5 — No correctly grouped terms found (total_terms={total_terms}, correct={correct_groupings})")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
