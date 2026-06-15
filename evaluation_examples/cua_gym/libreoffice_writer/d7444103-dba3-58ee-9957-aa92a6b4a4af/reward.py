"""
Reward Script: Create a complex invoice-style table in full_invoice.docx
Task ID: writer_tbl_075
Domain: libreoffice_writer
Scoring:
  Component 1: Table exists with 9 rows and 5 columns (0.15 pts)
  Component 2: Row 1 merged (A1-E1), 'TechCorp Solutions' bold 16pt centered (0.25 pts)
  Component 3: Row 2 merged (A2-E2), 'Invoice Date: 2024-03-15' (0.10 pts)
  Component 4: Row 3 header cells bold with gray background (0.10 pts)
  Component 5: 3 data rows correct values (0.10 pts)
  Component 6: Rows 7-8 merged labels (A-D) right-aligned with correct values (0.10 pts)
  Component 7: Row 9 TOTAL merged label bold right-aligned, E9=7150 (0.20 pts)
Total: 1.0
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'full_invoice'


def get_unique_cell(table, row_idx, col_idx):
    """Return cell at position, accounting for merged cells."""
    return table.cell(row_idx, col_idx)


def get_gridspan(cell):
    """Return gridSpan value of a cell (how many columns it spans)."""
    tc_pr = cell._tc.find(qn('w:tcPr'))
    if tc_pr is not None:
        gs = tc_pr.find(qn('w:gridSpan'))
        if gs is not None:
            val = gs.get(qn('w:val'))
            if val:
                return int(val)
    return 1


def get_shd_fill(cell):
    """Return background fill color of a cell, or None."""
    tc_pr = cell._tc.find(qn('w:tcPr'))
    if tc_pr is not None:
        shd = tc_pr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            return fill
    return None


def count_distinct_cells(row):
    """Count distinct (non-duplicated by merge) cells in a row."""
    return len(set(id(c) for c in row.cells))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No table found in document — document is empty or has no tables")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    t = doc.tables[0]

    # Component 1: Table has 9 rows and 5 columns (0.15 pts)
    try:
        row_count = len(t.rows)
        col_count = len(t.columns)
        if row_count == 9 and col_count == 5:
            print(f"PASS: Component 1 — Table has 9 rows and 5 columns (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Expected 9 rows x 5 cols, found {row_count} rows x {col_count} cols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row 1 (index 0) merged A1-E1, 'TechCorp Solutions', bold, 16pt, centered (0.25 pts)
    try:
        row0_distinct = count_distinct_cells(t.rows[0])
        cell_r0 = t.cell(0, 0)
        text_r0 = cell_r0.text.strip()

        # Check merge: row 0 should have only 1 distinct cell
        is_merged = (row0_distinct == 1)
        # Check text
        has_text = ('TechCorp Solutions' in text_r0)
        # Check bold
        is_bold = False
        for para in cell_r0.paragraphs:
            for run in para.runs:
                if run.font.bold is True or run.bold is True:
                    is_bold = True
        # Check font size (16pt)
        has_16pt = False
        for para in cell_r0.paragraphs:
            for run in para.runs:
                if run.font.size and abs(run.font.size.pt - 16.0) < 0.5:
                    has_16pt = True
        # Check centered alignment
        is_centered = False
        for para in cell_r0.paragraphs:
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                is_centered = True

        if is_merged and has_text and is_bold and has_16pt and is_centered:
            print(f"PASS: Component 2 — Row 1 merged A1-E1, 'TechCorp Solutions' bold 16pt centered (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — merged={is_merged}, text={has_text}, bold={is_bold}, 16pt={has_16pt}, centered={is_centered}; text='{text_r0}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Row 2 (index 1) merged A2-E2, 'Invoice Date: 2024-03-15' (0.10 pts)
    try:
        row1_distinct = count_distinct_cells(t.rows[1])
        cell_r1 = t.cell(1, 0)
        text_r1 = cell_r1.text.strip()

        is_merged = (row1_distinct == 1)
        has_text = ('Invoice Date: 2024-03-15' in text_r1)

        if is_merged and has_text:
            print(f"PASS: Component 3 — Row 2 merged A2-E2, 'Invoice Date: 2024-03-15' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — merged={is_merged}, text={has_text}; text='{text_r1}'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Row 3 (index 2) header cells bold with gray background (0.10 pts)
    try:
        expected_headers = ['Item', 'Desc', 'Qty', 'Rate', 'Amount']
        headers_correct = True
        headers_bold = True
        has_gray_bg = False

        for j, expected in enumerate(expected_headers):
            cell = t.cell(2, j)
            if cell.text.strip() != expected:
                headers_correct = False
                print(f"  Header col {j}: expected '{expected}', got '{cell.text.strip()}'")
            # Check bold
            for para in cell.paragraphs:
                for run in para.runs:
                    if not (run.font.bold is True or run.bold is True):
                        headers_bold = False

        # Check gray background on at least one header cell
        fill = get_shd_fill(t.cell(2, 0))
        if fill and fill.upper() not in ('', 'FFFFFF', 'AUTO', 'NONE', None):
            has_gray_bg = True

        if headers_correct and headers_bold and has_gray_bg:
            print(f"PASS: Component 4 — Row 3 headers correct, bold, gray background (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — headers_correct={headers_correct}, headers_bold={headers_bold}, gray_bg={has_gray_bg} (fill={fill})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 3 data rows (rows 4-6, indices 3-5) with correct values (0.10 pts)
    try:
        expected_data = [
            ['1', 'Web Development', '40', '100', '4000'],
            ['2', 'UI Design', '20', '80', '1600'],
            ['3', 'Testing', '15', '60', '900'],
        ]
        data_correct = True
        for row_i, expected_row in enumerate(expected_data):
            actual_row_idx = 3 + row_i
            for col_i, expected_val in enumerate(expected_row):
                actual_val = t.cell(actual_row_idx, col_i).text.strip()
                if actual_val != expected_val:
                    data_correct = False
                    print(f"  Data row {row_i+1} col {col_i}: expected '{expected_val}', got '{actual_val}'")

        if data_correct:
            print(f"PASS: Component 5 — 3 data rows with correct values (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — One or more data cells have incorrect values")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Rows 7-8 (indices 6-7) merged label cells A-D right-aligned, correct amounts (0.10 pts)
    try:
        # Row 7: A7-D7 merged, 'Subtotal:', right-aligned; E7='6500'
        # Row 8: A8-D8 merged, 'Tax (10%):', right-aligned; E8='650'
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        row6_distinct = count_distinct_cells(t.rows[6])
        row7_distinct = count_distinct_cells(t.rows[7])
        cell_r6_left = t.cell(6, 0)
        cell_r6_right = t.cell(6, 4)
        cell_r7_left = t.cell(7, 0)
        cell_r7_right = t.cell(7, 4)

        r6_merged = (row6_distinct == 2)
        r7_merged = (row7_distinct == 2)
        r6_label = ('Subtotal:' in cell_r6_left.text.strip())
        r7_label = ('Tax (10%):' in cell_r7_left.text.strip())
        r6_value = (cell_r6_right.text.strip() == '6500')
        r7_value = (cell_r7_right.text.strip() == '650')

        # Check right alignment
        r6_right_align = False
        for para in cell_r6_left.paragraphs:
            if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                r6_right_align = True

        r7_right_align = False
        for para in cell_r7_left.paragraphs:
            if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                r7_right_align = True

        if r6_merged and r7_merged and r6_label and r7_label and r6_value and r7_value and r6_right_align and r7_right_align:
            print(f"PASS: Component 6 — Rows 7-8 merged labels right-aligned with correct amounts (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — r6_merged={r6_merged}, r7_merged={r7_merged}, "
                  f"r6_label={r6_label}, r7_label={r7_label}, "
                  f"r6_value={r6_value} (got '{cell_r6_right.text.strip()}'), "
                  f"r7_value={r7_value} (got '{cell_r7_right.text.strip()}'), "
                  f"r6_align={r6_right_align}, r7_align={r7_right_align}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Row 9 (index 8) TOTAL merged label bold right-aligned, E9=7150 (0.20 pts)
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        row8_distinct = count_distinct_cells(t.rows[8])
        cell_r8_left = t.cell(8, 0)
        cell_r8_right = t.cell(8, 4)

        r8_merged = (row8_distinct == 2)
        r8_label_text = cell_r8_left.text.strip()
        r8_has_total = ('TOTAL:' in r8_label_text)
        r8_value = (cell_r8_right.text.strip() == '7150')

        r8_right_align = False
        r8_bold = False
        for para in cell_r8_left.paragraphs:
            if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                r8_right_align = True
            for run in para.runs:
                if run.font.bold is True or run.bold is True:
                    r8_bold = True

        if r8_merged and r8_has_total and r8_value and r8_right_align and r8_bold:
            print(f"PASS: Component 7 — Row 9 TOTAL merged bold right-aligned, E9=7150 (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 7 — r8_merged={r8_merged}, has_total={r8_has_total} ('{r8_label_text}'), "
                  f"value={r8_value} (got '{cell_r8_right.text.strip()}'), "
                  f"right_align={r8_right_align}, bold={r8_bold}")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
