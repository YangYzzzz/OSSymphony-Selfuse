"""
Initial Setup: Set column widths in a data table
Task ID: writer_tm_014
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_014'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Quarterly Sales Data Summary', level=1)

    # Add introductory paragraph
    intro = doc.add_paragraph(
        'The following table summarizes the regional sales performance '
        'for Q1 2025. All figures are in thousands of USD.'
    )

    # Create 4x10 table (header + 9 data rows) with default equal column widths
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Table Grid'

    # Headers
    headers = ['Region', 'Q1 Revenue', 'Units Sold', 'Avg Price']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Realistic data rows
    data = [
        ['Northeast',    '$127,450', '3,241', '$39.30'],
        ['Southeast',    '$98,230',  '2,876', '$34.15'],
        ['Midwest',      '$112,670', '3,102', '$36.32'],
        ['Southwest',    '$85,940',  '2,418', '$35.54'],
        ['West Coast',   '$143,820', '3,894', '$36.93'],
        ['Pacific NW',   '$67,310',  '1,987', '$33.87'],
        ['Mountain',     '$54,890',  '1,543', '$35.57'],
        ['Great Lakes',  '$91,750',  '2,631', '$34.88'],
        ['Mid-Atlantic', '$108,340', '2,975', '$36.42'],
    ]

    for i, row_data in enumerate(data, 1):
        for j, val in enumerate(row_data):
            table.cell(i, j).text = val

    # Add a closing paragraph
    doc.add_paragraph('')
    doc.add_paragraph(
        'Note: Data compiled from regional offices as of March 31, 2025. '
        'Figures subject to final audit adjustments.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
