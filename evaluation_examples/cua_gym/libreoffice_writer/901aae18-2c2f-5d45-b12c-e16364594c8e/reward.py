"""
Reward Script: Set column widths in a data table
Task ID: writer_tm_014
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Column 1 width is approximately 3cm
  Component 2 (0.45): Columns 2-4 width is approximately 4.5cm each
  Component 3 (0.20): Table structure and content preserved
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_014'

# Conversion: 1 cm = 567.0 twips (dxa units)
# 3 cm   = 1701 dxa
# 4.5 cm = 2551.5 dxa
# We allow 5% tolerance on widths

CM_TO_DXA = 567.0
TARGET_COL0_CM = 3.0
TARGET_COL123_CM = 4.5
TOLERANCE = 0.05  # 5%

TARGET_COL0_DXA = TARGET_COL0_CM * CM_TO_DXA   # 1701
TARGET_COL123_DXA = TARGET_COL123_CM * CM_TO_DXA  # 2551.5

# Expected table content (row texts for integrity check)
EXPECTED_ROW_COUNT = 10
EXPECTED_COL_COUNT = 4
EXPECTED_HEADERS = ['Region', 'Q1 Revenue', 'Units Sold', 'Avg Price']


def get_cell_width_dxa(cell):
    """Extract cell width in dxa (twips) from the XML tcPr/tcW element."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    w_elem = tc_pr.find(f'{{{ns}}}tcW')
    if w_elem is None:
        return None
    w_type = w_elem.get(f'{{{ns}}}type')
    w_val = int(w_elem.get(f'{{{ns}}}w'))
    if w_type == 'dxa':
        return w_val
    # If type is 'pct' or 'auto', we can't use it directly
    return None


def within_tolerance(actual, expected, tol=TOLERANCE):
    """Check if actual is within tol fraction of expected."""
    return abs(actual - expected) <= expected * tol


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Column 1 (index 0) width is approximately 3cm (0.35 points)
    try:
        # Check col 0 width across all rows
        col0_widths = []
        for row in table.rows:
            w = get_cell_width_dxa(row.cells[0])
            if w is not None:
                col0_widths.append(w)

        if len(col0_widths) == 0:
            print("FAIL: Component 1 — Could not read any column 0 widths")
        else:
            # All rows should have consistent width near 3cm
            all_match = all(within_tolerance(w, TARGET_COL0_DXA) for w in col0_widths)
            avg_width = sum(col0_widths) / len(col0_widths)
            avg_cm = avg_width / CM_TO_DXA

            if all_match:
                print(f"PASS: Component 1 — Column 1 width ~{avg_cm:.2f}cm (target 3.00cm) across {len(col0_widths)} rows (0.35 pts)")
                total_score += 0.35
            else:
                # Check if at least the average is correct (partial: not awarded here)
                print(f"FAIL: Component 1 — Column 1 avg width {avg_cm:.2f}cm, expected ~3.00cm. Widths: {col0_widths[:3]}...")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Columns 2-4 (indices 1,2,3) width is approximately 4.5cm each (0.45 points)
    try:
        col_pass_count = 0
        for col_idx in [1, 2, 3]:
            col_widths = []
            for row in table.rows:
                if col_idx < len(row.cells):
                    w = get_cell_width_dxa(row.cells[col_idx])
                    if w is not None:
                        col_widths.append(w)

            if len(col_widths) == 0:
                print(f"FAIL: Component 2 — Could not read column {col_idx+1} widths")
                continue

            all_match = all(within_tolerance(w, TARGET_COL123_DXA) for w in col_widths)
            avg_width = sum(col_widths) / len(col_widths)
            avg_cm = avg_width / CM_TO_DXA

            if all_match:
                print(f"PASS: Component 2.{col_idx} — Column {col_idx+1} width ~{avg_cm:.2f}cm (target 4.50cm)")
                col_pass_count += 1
            else:
                print(f"FAIL: Component 2.{col_idx} — Column {col_idx+1} avg width {avg_cm:.2f}cm, expected ~4.50cm. Widths: {col_widths[:3]}...")

        # Award proportional credit: 0.15 per column
        points = col_pass_count * 0.15
        if col_pass_count == 3:
            print(f"PASS: Component 2 — All 3 columns correct (0.45 pts)")
            total_score += 0.45
        elif col_pass_count > 0:
            print(f"PARTIAL: Component 2 — {col_pass_count}/3 columns correct ({points:.2f} pts)")
            total_score += points
        else:
            print("FAIL: Component 2 — No columns at target width")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table structure and content preserved (0.20 points)
    # This checks that the column width change didn't corrupt the table.
    # On initial_env the widths are wrong so components 1+2 fail (0.0),
    # but content IS preserved. However, we anchor this to a compound check:
    # content preserved AND col0 width differs from the equal-distribution default.
    # The equal default is 2160 dxa per column. If col0 is still 2160, this fails.
    try:
        row_count = len(table.rows)
        col_count = len(table.columns)

        # Sub-check: structure correct
        structure_ok = (row_count == EXPECTED_ROW_COUNT and col_count == EXPECTED_COL_COUNT)

        # Sub-check: headers correct
        headers = [table.cell(0, c).text.strip() for c in range(min(col_count, 4))]
        headers_ok = (headers == EXPECTED_HEADERS)

        # Sub-check: col0 width is NOT the default equal width (2160 dxa)
        # This ensures we only award points when widths have actually been changed
        col0_first = get_cell_width_dxa(table.rows[0].cells[0])
        width_changed = (col0_first is not None and not within_tolerance(col0_first, 2160, 0.02))

        if structure_ok and headers_ok and width_changed:
            print(f"PASS: Component 3 — Table preserved: {row_count} rows, {col_count} cols, headers match, widths changed (0.20 pts)")
            total_score += 0.20
        else:
            reasons = []
            if not structure_ok:
                reasons.append(f"structure {row_count}x{col_count} != {EXPECTED_ROW_COUNT}x{EXPECTED_COL_COUNT}")
            if not headers_ok:
                reasons.append(f"headers {headers} != {EXPECTED_HEADERS}")
            if not width_changed:
                reasons.append(f"col0 width {col0_first} still at default ~2160 (not changed)")
            print(f"FAIL: Component 3 — {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state(domain):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
