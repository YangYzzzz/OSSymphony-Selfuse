"""
Reward Script: Mail merge template setup in Welcome_Letter.docx
Task ID: writer_pd_019
Domain: libreoffice_writer
Scoring:
  - Component 1-5: Each merge field present (<<FirstName>>, <<LastName>>,
    <<Department>>, <<StartDate>>, <<ManagerName>>) — 0.15 pts each (0.75 total)
  - Component 6: All original placeholders removed — 0.15 pts
  - Component 7: Letter formatting preserved (paragraph structure intact) — 0.10 pts
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_019'

# The 5 merge fields that must be present, mapped from original placeholders
MERGE_FIELDS = {
    'FirstName': '[FIRST_NAME]',
    'LastName': '[LAST_NAME]',
    'Department': '[DEPARTMENT]',
    'StartDate': '[START_DATE]',
    'ManagerName': '[MANAGER_NAME]',
}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather full text from all paragraphs
    full_text = '\n'.join(p.text for p in doc.paragraphs)

    # Components 1-5: Check each merge field is present (0.15 pts each)
    for field_name, old_placeholder in MERGE_FIELDS.items():
        merge_tag = f'<<{field_name}>>'
        try:
            if merge_tag in full_text:
                print(f"PASS: Merge field {merge_tag} found (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Merge field {merge_tag} not found in document text")
        except Exception as e:
            print(f"ERROR: Could not check merge field {merge_tag}: {e}")

    # Component 6: All original placeholders removed (0.15 pts)
    try:
        remaining_placeholders = re.findall(r'\[(?:FIRST_NAME|LAST_NAME|DEPARTMENT|START_DATE|MANAGER_NAME)\]', full_text)
        if len(remaining_placeholders) == 0:
            print(f"PASS: All original placeholders removed (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: {len(remaining_placeholders)} original placeholder(s) still present: {remaining_placeholders}")
    except Exception as e:
        print(f"ERROR: Could not check placeholders: {e}")

    # Component 7: Letter formatting preserved AND merge fields present (0.10 pts)
    # This is a compound check: formatting intact + at least one merge field exists.
    # The merge field sub-condition ensures this only passes on golden (not initial).
    try:
        num_paras = len(doc.paragraphs)
        merge_field_count = len(re.findall(r'<<(?:FirstName|LastName|Department|StartDate|ManagerName)>>', full_text))
        if 11 <= num_paras <= 15 and merge_field_count >= 1:
            has_greeting = any('Dear' in p.text for p in doc.paragraphs)
            has_closing = any('Warm regards' in p.text for p in doc.paragraphs)
            has_company = any('Meridian Technologies' in p.text for p in doc.paragraphs)
            if has_greeting and has_closing and has_company:
                print(f"PASS: Letter formatting preserved with merge fields — {num_paras} paragraphs, {merge_field_count} fields (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Letter structure damaged — greeting={has_greeting}, closing={has_closing}, company={has_company}")
        else:
            print(f"FAIL: Formatting check — paragraphs={num_paras}, merge_fields={merge_field_count}")
    except Exception as e:
        print(f"ERROR: Could not check formatting: {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
