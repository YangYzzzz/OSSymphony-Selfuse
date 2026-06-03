"""
Initial Setup: Mail merge template with placeholder text
Task ID: writer_pd_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Company Header ---
    header_para = doc.add_paragraph()
    header_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.paragraph_format.space_after = Pt(4)
    run = header_para.add_run("Meridian Technologies Inc.")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_header = doc.add_paragraph()
    sub_header.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_header.paragraph_format.space_after = Pt(2)
    run = sub_header.add_run("Human Resources Department")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    addr = doc.add_paragraph()
    addr.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr.paragraph_format.space_after = Pt(18)
    run = addr.add_run("2500 Innovation Boulevard, Suite 400 | San Jose, CA 95134")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.add_run("March 15, 2026")
    run.font.size = Pt(11)

    # --- Salutation ---
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(12)
    run = salutation.add_run("Dear [FIRST_NAME] [LAST_NAME],")
    run.font.size = Pt(11)

    # --- Body Paragraph 1 ---
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(12)
    body1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body1.add_run(
        "On behalf of everyone at Meridian Technologies Inc., I am delighted to welcome you "
        "to our team. We are thrilled to have you join the "
    )
    run.font.size = Pt(11)
    run = body1.add_run("[DEPARTMENT]")
    run.font.size = Pt(11)
    run = body1.add_run(
        " department. Your skills and experience will be a tremendous asset to our organization, "
        "and we look forward to the contributions you will make."
    )
    run.font.size = Pt(11)

    # --- Body Paragraph 2 ---
    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(12)
    body2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body2.add_run(
        "Your official start date is "
    )
    run.font.size = Pt(11)
    run = body2.add_run("[START_DATE]")
    run.font.size = Pt(11)
    run = body2.add_run(
        ". Please arrive at the main reception desk by 9:00 AM, where a member of our "
        "onboarding team will guide you through the check-in process. You will receive your "
        "employee badge, workstation setup details, and a welcome kit during your first morning."
    )
    run.font.size = Pt(11)

    # --- Body Paragraph 3 ---
    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(12)
    body3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body3.add_run(
        "During your first week, you will participate in our comprehensive orientation program, "
        "which covers company policies, benefits enrollment, IT security training, and team "
        "introductions. Your direct manager, "
    )
    run.font.size = Pt(11)
    run = body3.add_run("[MANAGER_NAME]")
    run.font.size = Pt(11)
    run = body3.add_run(
        ", will meet with you on your first day to discuss your role, expectations, and "
        "initial projects. Please do not hesitate to reach out to them with any questions "
        "before your start date."
    )
    run.font.size = Pt(11)

    # --- Body Paragraph 4 ---
    body4 = doc.add_paragraph()
    body4.paragraph_format.space_after = Pt(12)
    body4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body4.add_run(
        "We believe that our people are our greatest strength, and we are confident that you "
        "will find Meridian Technologies to be a rewarding and supportive place to build your "
        "career. Should you need any assistance prior to your start date, please contact the "
        "HR department at hr@meridiantech.com or call (408) 555-0192."
    )
    run.font.size = Pt(11)

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(24)
    closing.paragraph_format.space_after = Pt(4)
    run = closing.add_run("Warm regards,")
    run.font.size = Pt(11)

    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_after = Pt(2)
    run = sig_name.add_run("Patricia Vasquez")
    run.bold = True
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(2)
    run = sig_title.add_run("Vice President, Human Resources")
    run.font.size = Pt(11)

    sig_company = doc.add_paragraph()
    run = sig_company.add_run("Meridian Technologies Inc.")
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
