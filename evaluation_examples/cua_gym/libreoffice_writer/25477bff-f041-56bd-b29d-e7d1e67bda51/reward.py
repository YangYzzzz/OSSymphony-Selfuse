"""
Reward Script: Modify list style indent for second-level items
Task ID: writer_lec_021
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Hanging indent of level 2 items ~= 567 twips (1.0 cm)
  Component 2 (0.4): Bullet/number position of level 2 items ~= 1134 twips (2.0 cm)
  Component 3 (0.2): Document text content is preserved
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_021'

# Conversion: 1 cm = 567 twips (360000 EMU / 635 = 566.93)
TWIPS_PER_CM = 567
TOLERANCE_TWIPS = 30  # ~0.05 cm tolerance

EXPECTED_HANGING = 1 * TWIPS_PER_CM   # 567 twips = 1.0 cm
EXPECTED_BULLET_POS = 2 * TWIPS_PER_CM  # 1134 twips = 2.0 cm

# Expected level-2 paragraph texts (for content integrity check)
EXPECTED_LEVEL2_TEXTS = [
    'Conduct competitive landscape review for Q3 targets',
    'Survey 200 enterprise customers on feature priorities',
    'Analyze pricing sensitivity across regional segments',
    'Complete API integration with Salesforce by July 15',
    'Finalize dashboard redesign with UX team approval',
    'Run load testing for 10,000 concurrent sessions',
    'Address critical security audit findings from PenTest report',
    'Draft press release and coordinate with PR agency',
    'Prepare demo environment for partner showcase events',
    'Update product landing page with new feature highlights',
    'Create updated battle cards for sales team distribution',
    'Schedule training sessions for EMEA and APAC regions',
    'Set up real-time analytics dashboard for adoption metrics',
    'Establish customer feedback loop with support escalation path',
    'Plan 30-day retrospective with all stakeholders',
]


def persist_app_state(domain):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_level1_indent_from_numbering(doc):
    """
    Extract the indent properties (left, hanging) for ilvl=1 of the
    abstractNum that is referenced by numId=10 in the numbering part.
    Also checks per-paragraph overrides on ilvl=1 paragraphs.
    Returns (left_twips, hanging_twips) from the numbering definition,
    or from per-paragraph overrides if they exist.
    """
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    wns = ns['w']

    # First try numbering definition
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return None, None

    root = numbering_part.element

    # Find which abstractNumId is mapped to numId used by the list
    # We need to find the numId used by level-2 paragraphs
    target_num_id = None
    for para in doc.paragraphs:
        numPr = para._element.find('.//w:numPr', ns)
        if numPr is not None:
            ilvl_el = numPr.find('w:ilvl', ns)
            numId_el = numPr.find('w:numId', ns)
            if ilvl_el is not None and numId_el is not None:
                ilvl_val = ilvl_el.get(f'{{{wns}}}val')
                if ilvl_val == '1':
                    target_num_id = numId_el.get(f'{{{wns}}}val')
                    break

    if target_num_id is None:
        print("WARN: No level-2 paragraphs found in document")
        return None, None

    # Map numId -> abstractNumId
    abstract_num_id = None
    for num in root.findall('.//w:num', ns):
        nid = num.get(f'{{{wns}}}numId')
        if nid == target_num_id:
            aref = num.find('w:abstractNumId', ns)
            if aref is not None:
                abstract_num_id = aref.get(f'{{{wns}}}val')
            break

    if abstract_num_id is None:
        print(f"WARN: Could not find abstractNumId for numId={target_num_id}")
        return None, None

    # Find the abstractNum and get level 1 indent
    for abstract_num in root.findall('.//w:abstractNum', ns):
        abs_id = abstract_num.get(f'{{{wns}}}abstractNumId')
        if abs_id == abstract_num_id:
            for lvl in abstract_num.findall('.//w:lvl', ns):
                ilvl = lvl.get(f'{{{wns}}}ilvl')
                if ilvl == '1':
                    ppr = lvl.find('.//w:pPr', ns)
                    if ppr is not None:
                        ind = ppr.find('w:ind', ns)
                        if ind is not None:
                            left = ind.get(f'{{{wns}}}left')
                            hanging = ind.get(f'{{{wns}}}hanging')
                            left_val = int(left) if left else None
                            hanging_val = int(hanging) if hanging else None
                            print(f"INFO: Numbering definition level 1: left={left_val}, hanging={hanging_val}")
                            return left_val, hanging_val

    # Also check for per-paragraph indent overrides (some implementations do this)
    for para in doc.paragraphs:
        numPr = para._element.find('.//w:numPr', ns)
        if numPr is not None:
            ilvl_el = numPr.find('w:ilvl', ns)
            if ilvl_el is not None and ilvl_el.get(f'{{{wns}}}val') == '1':
                pPr = para._element.find('w:pPr', ns)
                if pPr is not None:
                    ind = pPr.find('w:ind', ns)
                    if ind is not None:
                        left = ind.get(f'{{{wns}}}left')
                        hanging = ind.get(f'{{{wns}}}hanging')
                        if left and hanging:
                            left_val = int(left)
                            hanging_val = int(hanging)
                            print(f"INFO: Per-paragraph override level 1: left={left_val}, hanging={hanging_val}")
                            return left_val, hanging_val

    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get level 1 (second-level) indent properties
    left_twips, hanging_twips = get_level1_indent_from_numbering(doc)

    if left_twips is None or hanging_twips is None:
        print("FAIL: Could not extract level 1 indent from numbering definition")
        print("REWARD: 0.0")
        return 0.0

    bullet_pos = left_twips - hanging_twips

    # Component 1: Hanging indent ~= 567 twips (1.0 cm) — 0.4 points
    try:
        diff_hanging = abs(hanging_twips - EXPECTED_HANGING)
        if diff_hanging <= TOLERANCE_TWIPS:
            print(f"PASS: Component 1 — hanging indent = {hanging_twips} twips "
                  f"(expected ~{EXPECTED_HANGING}, diff={diff_hanging}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — hanging indent = {hanging_twips} twips "
                  f"(expected ~{EXPECTED_HANGING}, diff={diff_hanging})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Bullet position (left - hanging) ~= 1134 twips (2.0 cm) — 0.4 points
    try:
        diff_bullet = abs(bullet_pos - EXPECTED_BULLET_POS)
        if diff_bullet <= TOLERANCE_TWIPS:
            print(f"PASS: Component 2 — bullet position = {bullet_pos} twips "
                  f"(expected ~{EXPECTED_BULLET_POS}, diff={diff_bullet}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — bullet position = {bullet_pos} twips "
                  f"(expected ~{EXPECTED_BULLET_POS}, diff={diff_bullet})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document text content is preserved — 0.2 points
    # Only scores if the task change (components 1 or 2) also passed,
    # to avoid awarding points on unchanged initial_env.
    try:
        if total_score > 0:
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            wns = ns['w']
            level2_texts = []
            for para in doc.paragraphs:
                numPr = para._element.find('.//w:numPr', ns)
                if numPr is not None:
                    ilvl_el = numPr.find('w:ilvl', ns)
                    if ilvl_el is not None and ilvl_el.get(f'{{{wns}}}val') == '1':
                        level2_texts.append(para.text.strip())

            if level2_texts == [t.strip() for t in EXPECTED_LEVEL2_TEXTS]:
                print(f"PASS: Component 3 — all {len(level2_texts)} level-2 texts preserved (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — level-2 texts changed or missing "
                      f"(found {len(level2_texts)}, expected {len(EXPECTED_LEVEL2_TEXTS)})")
        else:
            print("SKIP: Component 3 — no task-change detected, skipping content check")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
