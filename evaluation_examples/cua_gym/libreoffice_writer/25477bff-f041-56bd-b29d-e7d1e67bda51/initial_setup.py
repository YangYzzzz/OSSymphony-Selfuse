"""
Initial Setup: Create a Writer document with a two-level bulleted list using default position settings.
Task ID: writer_lec_021
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_021'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    title = doc.add_heading('Project Planning Checklist', level=1)

    # Add an introductory paragraph
    intro = doc.add_paragraph(
        'The following checklist outlines the key areas and action items '
        'for the Q3 product launch initiative.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # --- Build a custom two-level bullet list via numbering XML ---
    # We need to define a numbering definition with two levels.
    # Level 0: bullet at default (0.63 cm indent, 1.27 cm text)
    # Level 1: bullet at default (1.27 cm indent, 1.9 cm text) -- DEFAULT positions

    # Access the numbering part (create if needed)
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part._element

    # Define an abstract numbering with two bullet levels
    # Default positions: level 0 at ~0.63cm, level 1 at ~1.27cm
    abstract_num_xml = f'''
    <w:abstractNum w:abstractNumId="10" {nsdecls('w')}>
        <w:multiLevelType w:val="hybridMultilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="\u2022"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="\u25CB"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="1440" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:hint="default"/>
            </w:rPr>
        </w:lvl>
    </w:abstractNum>
    '''
    abstract_num = parse_xml(abstract_num_xml)
    numbering_elem.append(abstract_num)

    # Create a concrete num referencing this abstractNum
    num_xml = f'''
    <w:num w:numId="10" {nsdecls('w')}>
        <w:abstractNumId w:val="10"/>
    </w:num>
    '''
    num_elem = parse_xml(num_xml)
    numbering_elem.append(num_elem)

    # Helper to add a bullet paragraph at a given level
    def add_bullet(text, level=0):
        para = doc.add_paragraph()
        para.text = text
        para.style = doc.styles['Normal']
        # Set numbering reference
        pPr = para._element.get_or_add_pPr()
        numPr = parse_xml(f'<w:numPr {nsdecls("w")}><w:ilvl w:val="{level}"/><w:numId w:val="10"/></w:numPr>')
        pPr.append(numPr)
        return para

    # --- Content: Project planning checklist ---
    # Level 0 items with level 1 sub-items

    add_bullet('Market Research and Analysis', level=0)
    add_bullet('Conduct competitive landscape review for Q3 targets', level=1)
    add_bullet('Survey 200 enterprise customers on feature priorities', level=1)
    add_bullet('Analyze pricing sensitivity across regional segments', level=1)

    add_bullet('Product Development Milestones', level=0)
    add_bullet('Complete API integration with Salesforce by July 15', level=1)
    add_bullet('Finalize dashboard redesign with UX team approval', level=1)
    add_bullet('Run load testing for 10,000 concurrent sessions', level=1)
    add_bullet('Address critical security audit findings from PenTest report', level=1)

    add_bullet('Marketing and Go-to-Market Strategy', level=0)
    add_bullet('Draft press release and coordinate with PR agency', level=1)
    add_bullet('Prepare demo environment for partner showcase events', level=1)
    add_bullet('Update product landing page with new feature highlights', level=1)

    add_bullet('Sales Enablement', level=0)
    add_bullet('Create updated battle cards for sales team distribution', level=1)
    add_bullet('Schedule training sessions for EMEA and APAC regions', level=1)

    add_bullet('Post-Launch Monitoring', level=0)
    add_bullet('Set up real-time analytics dashboard for adoption metrics', level=1)
    add_bullet('Establish customer feedback loop with support escalation path', level=1)
    add_bullet('Plan 30-day retrospective with all stakeholders', level=1)

    # Add a closing paragraph
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.add_run('Note: ').bold = True
    closing.add_run(
        'All milestones should be reviewed bi-weekly during the steering committee meetings. '
        'Contact the PMO office for resource allocation questions.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
