"""
Reward Script: Add headers/footers to white paper
Task ID: writer_mktg_005
Domain: libreoffice_writer
Scoring:
  Component 1: different_first_page enabled (0.10 pts)
  Component 2: Header text 'Zenith Analytics' present (0.25 pts)
  Component 3: Header right-aligned, 9pt, gray (#808080) (0.25 pts)
  Component 4: First page header is empty (0.10 pts)
  Component 5: Footer confidential text present and italic 8pt (0.15 pts)
  Component 6: Footer contains PAGE/NUMPAGES field codes (0.15 pts)
  Total: 1.0
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'data_analytics_whitepaper'

FILE_PATH = f'{WORKDIR}/{TASK_ID}.docx'


def persist_app_state():
    """Send ctrl+s to save any unsaved GUI edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Add headers and footers to a white paper:
    - Header on pages 2-6: 'Zenith Analytics' right-aligned, 9pt, gray (#808080)
    - First page (title page): NO header (different_first_page must be enabled)
    - Footer on all pages: 'Confidential — For Authorized Distribution Only' on left (italic, 8pt)
    - Footer page number: 'Page X of Y' centered (using PAGE/NUMPAGES field codes)
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: different_first_page is enabled (0.10 pts)
    # This ensures title page (page 1) has no header while pages 2+ do.
    try:
        dfp = section.different_first_page_header_footer
        if dfp:
            print(f"PASS: Component 1 — different_first_page is enabled (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — different_first_page is False; expected True for title page exclusion")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header text 'Zenith Analytics' present in default (non-first-page) header (0.25 pts)
    try:
        header = section.header
        header_texts = [p.text.strip() for p in header.paragraphs]
        header_full = ' '.join(header_texts).strip()
        if 'Zenith Analytics' in header_full:
            print(f"PASS: Component 2 — Header contains 'Zenith Analytics' (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Header text '{header_full}' does not contain 'Zenith Analytics'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header formatting — right-aligned, 9pt, gray color #808080 (0.25 pts)
    try:
        header = section.header
        format_score = 0.0
        format_checks = []

        for p in header.paragraphs:
            if 'Zenith Analytics' not in p.text:
                continue
            # Check alignment
            alignment = p.paragraph_format.alignment
            if alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                format_checks.append('right-aligned')
            else:
                print(f"  DETAIL: Header alignment is {alignment}, expected RIGHT (2)")

            # Check runs for size and color
            for run in p.runs:
                if 'Zenith Analytics' not in run.text:
                    continue
                # Check size
                sz = run.font.size
                if sz and abs(sz.pt - 9.0) < 0.5:
                    format_checks.append('9pt')
                elif sz:
                    print(f"  DETAIL: Header font size is {sz.pt}pt, expected 9pt")
                else:
                    print(f"  DETAIL: Header font size is None (inherited)")

                # Check color — should be gray #808080 = RGB(128, 128, 128)
                try:
                    color = run.font.color.rgb
                    if color:
                        r, g, b = color[0], color[1], color[2]
                        # Accept gray range: each channel between 100-160
                        if abs(r - 128) <= 30 and abs(g - 128) <= 30 and abs(b - 128) <= 30 and abs(r - g) <= 20 and abs(g - b) <= 20:
                            format_checks.append('gray-color')
                        else:
                            print(f"  DETAIL: Header color RGB({r},{g},{b}), expected ~(128,128,128)")
                    else:
                        print(f"  DETAIL: Header color is None (not explicitly set)")
                except Exception as ce:
                    print(f"  DETAIL: color check error: {ce}")

        # Award points if at least 2 out of 3 formatting checks pass
        unique_checks = set(format_checks)
        passed = len(unique_checks)
        if 'right-aligned' in unique_checks and '9pt' in unique_checks and 'gray-color' in unique_checks:
            print(f"PASS: Component 3 — Header formatting correct: right-aligned, 9pt, gray (0.25 pts)")
            total_score += 0.25
        elif passed >= 2:
            print(f"PARTIAL: Component 3 — Header formatting partially correct ({passed}/3 checks: {unique_checks}) (0.12 pts)")
            total_score += 0.12
        else:
            print(f"FAIL: Component 3 — Header formatting incorrect ({passed}/3 checks: {unique_checks})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: First page header is empty AND different_first_page is enabled (0.10 pts)
    # This compound check ensures the title page is specifically excluded via the
    # different_first_page_header_footer mechanism, not just because headers are absent.
    # It FAILS on initial (different_first_page=False) and PASSES on golden (True + empty first page).
    try:
        dfp_enabled = section.different_first_page_header_footer
        first_hdr = section.first_page_header
        first_hdr_texts = [p.text.strip() for p in first_hdr.paragraphs]
        first_hdr_full = ' '.join(first_hdr_texts).strip()
        # Require BOTH: different_first_page=True AND first page header empty
        if dfp_enabled and not first_hdr_full:
            print(f"PASS: Component 4 — different_first_page enabled + first page header is empty (title page excluded) (0.10 pts)")
            total_score += 0.10
        elif not dfp_enabled:
            print(f"FAIL: Component 4 — different_first_page is False; title page not excluded from header")
        else:
            print(f"FAIL: Component 4 — First page header contains '{first_hdr_full}'; should be empty")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Footer contains 'Confidential — For Authorized Distribution Only' (italic, ~8pt) (0.15 pts)
    try:
        footer = section.footer
        confidential_text = 'Confidential'
        full_footer_text = ''
        found_confidential = False
        found_italic = False
        found_8pt = False

        for p in footer.paragraphs:
            full_footer_text += p.text
            for run in p.runs:
                if 'Confidential' in run.text:
                    found_confidential = True
                    # Check italic
                    if run.italic:
                        found_italic = True
                    # Check size ~8pt
                    sz = run.font.size
                    if sz and abs(sz.pt - 8.0) < 0.5:
                        found_8pt = True
                    elif sz:
                        print(f"  DETAIL: Confidential text size is {sz.pt}pt, expected 8pt")

        if found_confidential and found_italic:
            print(f"PASS: Component 5 — Footer has confidential text, italic (0.15 pts)")
            total_score += 0.15
            if not found_8pt:
                print(f"  NOTE: Confidential text size may not be exactly 8pt")
        elif found_confidential:
            print(f"PARTIAL: Component 5 — Footer has confidential text but not italic (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 5 — Footer text '{full_footer_text[:80]}' missing 'Confidential — For Authorized Distribution Only'")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Footer contains PAGE and NUMPAGES field codes (page number centered) (0.15 pts)
    try:
        footer = section.footer
        found_page_field = False
        found_numpages_field = False

        for p in footer.paragraphs:
            xml = p._element.xml
            instrs = re.findall(r'<w:instrText[^>]*>(.*?)</w:instrText>', xml)
            instr_str = ' '.join(instrs)
            if 'PAGE' in instr_str:
                found_page_field = True
            if 'NUMPAGES' in instr_str:
                found_numpages_field = True

        if found_page_field and found_numpages_field:
            print(f"PASS: Component 6 — Footer has PAGE and NUMPAGES field codes (0.15 pts)")
            total_score += 0.15
        elif found_page_field:
            print(f"PARTIAL: Component 6 — Footer has PAGE field but missing NUMPAGES (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 6 — Footer missing page number field codes")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved state (task involves LibreOffice Writer)
persist_app_state()

if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
