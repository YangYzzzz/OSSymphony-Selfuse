"""
Initial Setup: Data Analytics White Paper - No headers/footers
Task ID: writer_mktg_005
Domain: libreoffice_writer

Creates a 6-page white paper docx with title page, TOC, and 4 content pages.
No headers or footers (agent will add them).
File is placed at /home/user/Desktop/data_analytics_whitepaper.docx
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_005'
FILENAME = 'data_analytics_whitepaper.docx'
OUTPUT = f'{WORKDIR}/{FILENAME}'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add an explicit page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)
    return para


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # -----------------------------------------------------------------------
    # PAGE 1: Title Page
    # -----------------------------------------------------------------------
    # Spacer at top
    for _ in range(6):
        doc.add_paragraph()

    # Main title - centered, 20pt bold
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('Harnessing Predictive Analytics for Supply Chain Optimization')
    title_run.bold = True
    title_run.font.size = Pt(20)

    # Subtitle
    doc.add_paragraph()
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('A Strategic White Paper')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Author / org info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_run = author_para.add_run('Zenith Analytics Research Division')
    author_run.font.size = Pt(12)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run('March 2025')
    date_run.font.size = Pt(12)

    # Page break to end title page
    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 2: Table of Contents
    # -----------------------------------------------------------------------
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    toc_run = toc_heading.add_run('Table of Contents')
    toc_run.bold = True
    toc_run.font.size = Pt(16)

    doc.add_paragraph()

    toc_entries = [
        ('1.  Executive Summary', '3'),
        ('2.  The Challenge: Supply Chain Visibility Gaps', '3'),
        ('3.  Predictive Analytics: Core Capabilities', '4'),
        ('4.  Implementation Framework', '5'),
        ('5.  Case Study: GlobalManufacture Inc.', '5'),
        ('6.  Conclusion & Recommendations', '6'),
    ]

    for entry_text, page_num in toc_entries:
        toc_entry = doc.add_paragraph()
        toc_entry.paragraph_format.left_indent = Inches(0.25)
        entry_run = toc_entry.add_run(f'{entry_text}{"." * (50 - len(entry_text))}{page_num}')
        entry_run.font.size = Pt(11)

    # Page break to end TOC page
    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 3: Executive Summary & Challenge
    # -----------------------------------------------------------------------
    h1 = doc.add_paragraph()
    r = h1.add_run('1. Executive Summary')
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        'In today\'s interconnected global economy, supply chain disruptions cost enterprises '
        'an estimated $184 billion annually. Zenith Analytics has identified predictive analytics '
        'as the single most impactful technology investment for supply chain resilience. This white '
        'paper presents our findings from a 24-month study spanning 47 enterprise clients across '
        'manufacturing, retail, and logistics sectors.'
    ).runs[0].font.size = Pt(11)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    r2 = h2.add_run('2. The Challenge: Supply Chain Visibility Gaps')
    r2.bold = True
    r2.font.size = Pt(14)

    doc.add_paragraph(
        'Modern supply chains span dozens of countries, hundreds of suppliers, and thousands of '
        'SKUs. Traditional reactive approaches — responding to disruptions after they occur — are '
        'no longer sufficient. Key pain points identified in our research include:'
    ).runs[0].font.size = Pt(11)

    challenges = [
        'Limited real-time inventory visibility across multi-tier supplier networks',
        'Inability to correlate macroeconomic signals with operational risk',
        'Manual demand forecasting with 23% average error rates',
        'Siloed data systems preventing cross-functional decision making',
        'Delayed identification of supplier financial stress indicators',
    ]
    for challenge in challenges:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.add_run(challenge).font.size = Pt(11)

    doc.add_paragraph(
        'These visibility gaps result in excess inventory buffers averaging 34% above optimal '
        'levels, while simultaneously experiencing stockout events 2.8 times per quarter on '
        'critical components.'
    ).runs[0].font.size = Pt(11)

    # Page break
    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 4: Core Capabilities & Implementation
    # -----------------------------------------------------------------------
    h3 = doc.add_paragraph()
    r3 = h3.add_run('3. Predictive Analytics: Core Capabilities')
    r3.bold = True
    r3.font.size = Pt(14)

    doc.add_paragraph(
        'The Zenith Analytics Predictive Supply Chain Platform (PSCP) integrates machine learning '
        'models trained on over 2.3 billion historical transaction records. Our platform delivers '
        'three core capability clusters:'
    ).runs[0].font.size = Pt(11)

    # Sub-headings with content
    sub_sections = [
        ('3.1 Demand Sensing & Forecasting',
         'PSCP ingests 140+ external signals — including weather patterns, social media sentiment, '
         'competitor pricing, and economic indicators — to generate rolling 90-day demand forecasts '
         'with mean absolute percentage error (MAPE) of 7.2%, compared to the industry average of '
         '23%. Clients report inventory carrying cost reductions of 18-31%.'),
        ('3.2 Supplier Risk Intelligence',
         'Our NLP-driven supplier monitoring system continuously analyzes news feeds, financial '
         'filings, shipping data, and geopolitical events to generate risk scores for 450,000+ '
         'global suppliers. Average early warning lead time is 47 days before a supply disruption '
         'materializes.'),
        ('3.3 Dynamic Reorder Optimization',
         'Reinforcement learning algorithms continuously optimize reorder points and safety stock '
         'levels based on real-time demand signals, supplier lead time variability, and carrying '
         'costs. Clients achieve 99.1% service levels while reducing total inventory investment.'),
    ]

    for sub_title, sub_text in sub_sections:
        doc.add_paragraph()
        sh = doc.add_paragraph()
        sr = sh.add_run(sub_title)
        sr.bold = True
        sr.italic = True
        sr.font.size = Pt(12)
        doc.add_paragraph(sub_text).runs[0].font.size = Pt(11)

    # Page break
    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 5: Implementation Framework & Case Study
    # -----------------------------------------------------------------------
    h4 = doc.add_paragraph()
    r4 = h4.add_run('4. Implementation Framework')
    r4.bold = True
    r4.font.size = Pt(14)

    doc.add_paragraph(
        'Successful deployment of predictive analytics requires a structured approach. '
        'Our proven four-phase framework minimizes disruption while accelerating time-to-value:'
    ).runs[0].font.size = Pt(11)

    phases = [
        ('Phase 1 — Discovery (Weeks 1-4):', 'Data audit, systems integration mapping, KPI baseline establishment, stakeholder alignment workshops.'),
        ('Phase 2 — Foundation (Weeks 5-12):', 'Data pipeline construction, historical model training, integration with ERP/WMS systems, initial dashboard deployment.'),
        ('Phase 3 — Activation (Weeks 13-20):', 'Model validation, user training, parallel running with legacy systems, iterative refinement based on operational feedback.'),
        ('Phase 4 — Optimization (Week 21+):', 'Continuous model retraining, expansion to additional supplier tiers, advanced scenario planning, ROI measurement and reporting.'),
    ]

    for phase_title, phase_desc in phases:
        p = doc.add_paragraph(style='List Number')
        run_title = p.add_run(phase_title + ' ')
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_desc = p.add_run(phase_desc)
        run_desc.font.size = Pt(11)

    doc.add_paragraph()

    h5 = doc.add_paragraph()
    r5 = h5.add_run('5. Case Study: GlobalManufacture Inc.')
    r5.bold = True
    r5.font.size = Pt(14)

    doc.add_paragraph(
        'GlobalManufacture Inc. (GMI), a $3.2B industrial equipment manufacturer with operations '
        'in 18 countries, deployed PSCP in Q1 2024. Prior to implementation, GMI experienced '
        'chronic component shortages averaging 4.2 per month, resulting in production line '
        'stoppages costing approximately $2.1M per incident.'
    ).runs[0].font.size = Pt(11)

    doc.add_paragraph(
        'Results after 12 months of operation: Component shortage incidents reduced by 78% '
        '(from 4.2 to 0.9 per month), inventory carrying costs reduced by $14.3M annually, '
        'on-time delivery performance improved from 84.2% to 97.6%, and supplier risk events '
        'flagged with average 52-day advance notice.'
    ).runs[0].font.size = Pt(11)

    # Page break
    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 6: Conclusion
    # -----------------------------------------------------------------------
    h6 = doc.add_paragraph()
    r6 = h6.add_run('6. Conclusion & Recommendations')
    r6.bold = True
    r6.font.size = Pt(14)

    doc.add_paragraph(
        'Predictive analytics represents the most significant opportunity for supply chain '
        'transformation in the coming decade. Organizations that invest now in data infrastructure '
        'and analytical capabilities will achieve sustainable competitive advantages through '
        'superior resilience, cost efficiency, and customer service levels.'
    ).runs[0].font.size = Pt(11)

    doc.add_paragraph()

    recs_heading = doc.add_paragraph()
    rh_run = recs_heading.add_run('Key Recommendations:')
    rh_run.bold = True
    rh_run.font.size = Pt(12)

    recommendations = [
        'Initiate a supply chain data maturity assessment within the next 90 days',
        'Identify 2-3 high-impact use cases for initial predictive analytics deployment',
        'Establish cross-functional data governance committee with C-suite sponsorship',
        'Allocate dedicated budget for data integration and quality improvement programs',
        'Develop internal analytics talent alongside external platform partnerships',
    ]
    for rec in recommendations:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.add_run(rec).font.size = Pt(11)

    doc.add_paragraph()

    closing_para = doc.add_paragraph()
    closing_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    closing_run = closing_para.add_run(
        'For more information, contact Zenith Analytics at info@zenithanalytics.com\n'
        'or visit www.zenithanalytics.com/supply-chain'
    )
    closing_run.font.size = Pt(10)
    closing_run.italic = True

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
