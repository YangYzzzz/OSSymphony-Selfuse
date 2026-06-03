"""
Initial Setup: Set thesis document margins
Task ID: writer_acad_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: Default margins (2 cm all around) ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # --- Title Page ---
    for _ in range(4):
        doc.add_paragraph('')

    title = doc.add_heading('', level=0)
    run = title.add_run('Emergent Dynamics in Complex Adaptive Systems:\nA Multi-Scale Analysis of Self-Organization')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Dissertation Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy')
    run.font.size = Pt(14)

    doc.add_paragraph('')

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Elena Marchetti')
    run.font.size = Pt(16)
    run.bold = True

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run('Department of Applied Mathematics and Computational Science\nHarrington University\nMay 2025')
    run.font.size = Pt(12)

    doc.add_paragraph('')

    committee = doc.add_paragraph()
    committee.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = committee.add_run('Dissertation Committee:\nProf. David Nakamura (Chair)\nProf. Amira Okafor\nProf. Lars Johansson\nDr. Priya Sundaram')
    run.font.size = Pt(11)

    # --- Page break: Abstract ---
    doc.add_page_break()

    abstract_heading = doc.add_heading('Abstract', level=1)
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        'This dissertation investigates the emergent dynamics of complex adaptive systems '
        'through a multi-scale analytical framework that bridges microscopic agent-level interactions '
        'and macroscopic pattern formation. We develop a novel mathematical formalism that unifies '
        'concepts from statistical mechanics, network theory, and dynamical systems to characterize '
        'the conditions under which self-organized criticality arises in heterogeneous populations.'
    )

    doc.add_paragraph(
        'Our primary contribution is the introduction of the Adaptive Coupling Framework (ACF), '
        'which extends classical mean-field approximations by incorporating local feedback mechanisms '
        'and stochastic perturbations. Through extensive numerical simulations involving over 10^6 '
        'agent interactions across 500 independent trials, we demonstrate that ACF predictions '
        'achieve a correlation coefficient of r = 0.94 with observed emergent behavior in three '
        'distinct experimental domains: neural network synchronization, ant colony foraging optimization, '
        'and urban traffic flow dynamics.'
    )

    doc.add_paragraph(
        'We further establish theoretical bounds on the critical thresholds for phase transitions '
        'in adaptive networks, proving that the onset of global coherence is governed by a universal '
        'scaling law with exponent beta = 0.63 +/- 0.04. These results challenge existing models '
        'that assume homogeneous coupling strengths and provide a more nuanced understanding of '
        'how local adaptation drives system-wide reorganization.'
    )

    # --- Page break: Chapter 1 ---
    doc.add_page_break()

    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('1.1 Background and Motivation', level=2)

    doc.add_paragraph(
        'Complex adaptive systems (CAS) represent one of the most fascinating and challenging '
        'areas of modern scientific inquiry. From the intricate neural circuits that give rise to '
        'consciousness to the decentralized decision-making processes of social insect colonies, '
        'these systems exhibit a remarkable property: the emergence of sophisticated global behavior '
        'from relatively simple local interactions (Holland, 1995; Kauffman, 1993).'
    )

    doc.add_paragraph(
        'The study of emergence in complex systems has a rich intellectual history spanning multiple '
        'disciplines. Anderson\'s seminal 1972 paper "More Is Different" established the conceptual '
        'foundation by arguing that each level of organizational complexity introduces fundamentally '
        'new properties that cannot be predicted from lower-level descriptions alone. This insight '
        'has since been formalized through various mathematical frameworks, including synergetics '
        '(Haken, 1983), autopoiesis (Maturana & Varela, 1980), and the theory of dissipative '
        'structures (Prigogine & Stengers, 1984).'
    )

    doc.add_paragraph(
        'Despite these advances, a unified theoretical framework capable of predicting emergent '
        'behavior across different classes of complex systems remains elusive. Current models tend '
        'to be domain-specific, requiring substantial modification when applied to new systems. '
        'The mean-field approximation, while computationally tractable, often fails to capture the '
        'heterogeneous coupling dynamics that are characteristic of real-world adaptive systems '
        '(Bar-Yam, 2003; Mitchell, 2009).'
    )

    doc.add_heading('1.2 Research Questions', level=2)

    doc.add_paragraph(
        'This dissertation addresses three interconnected research questions that aim to advance '
        'our understanding of emergent dynamics in complex adaptive systems:'
    )

    doc.add_paragraph(
        'RQ1: Under what conditions does self-organized criticality emerge in heterogeneous '
        'populations with adaptive coupling strengths, and how do these conditions differ from '
        'predictions based on homogeneous mean-field models?',
        style='List Number'
    )

    doc.add_paragraph(
        'RQ2: Can a unified mathematical framework based on local feedback mechanisms and '
        'stochastic perturbations accurately predict phase transition thresholds across multiple '
        'classes of complex adaptive systems?',
        style='List Number'
    )

    doc.add_paragraph(
        'RQ3: What is the relationship between the topology of the interaction network and the '
        'universality class of the emergent phase transition, and how does adaptive rewiring '
        'influence this relationship?',
        style='List Number'
    )

    doc.add_heading('1.3 Scope and Contributions', level=2)

    doc.add_paragraph(
        'The primary contributions of this work are threefold. First, we introduce the Adaptive '
        'Coupling Framework (ACF), a mathematical formalism that extends classical mean-field '
        'theory by incorporating heterogeneous agent-level adaptation dynamics. Second, we provide '
        'rigorous numerical validation of ACF predictions across three experimental domains, '
        'demonstrating its generalizability. Third, we derive analytical bounds on critical '
        'thresholds for phase transitions that yield a universal scaling law.'
    )

    # --- Page break: Chapter 2 ---
    doc.add_page_break()

    doc.add_heading('Chapter 2: Literature Review', level=1)

    doc.add_heading('2.1 Foundations of Complexity Science', level=2)

    doc.add_paragraph(
        'The intellectual roots of complexity science can be traced to several convergent '
        'developments in the mid-twentieth century. Cybernetics (Wiener, 1948), general systems '
        'theory (von Bertalanffy, 1968), and information theory (Shannon, 1948) each contributed '
        'essential concepts that would later coalesce into the modern study of complex systems. '
        'The Santa Fe Institute, founded in 1984, played a pivotal role in establishing complexity '
        'as a legitimate interdisciplinary field of study (Waldrop, 1992).'
    )

    doc.add_paragraph(
        'A fundamental concept in complexity science is the notion of emergence, which refers to '
        'the appearance of system-level properties that are not present in any individual component. '
        'Bedau (1997) distinguished between "weak emergence," where macro-level properties can in '
        'principle be derived from micro-level descriptions through simulation, and "strong emergence," '
        'where such derivation is impossible even in principle. This dissertation primarily concerns '
        'itself with weak emergence, as our mathematical framework aims to predict emergent behavior '
        'from knowledge of agent-level interactions.'
    )

    doc.add_heading('2.2 Self-Organized Criticality', level=2)

    doc.add_paragraph(
        'Self-organized criticality (SOC), introduced by Bak, Tang, and Wiesenfeld (1987), '
        'describes the tendency of certain complex systems to naturally evolve toward a critical '
        'state characterized by power-law distributions and scale-invariant behavior. The canonical '
        'example is the sandpile model, where grains of sand dropped onto a pile produce avalanches '
        'whose sizes follow a power law distribution P(s) ~ s^(-tau) with tau approximately 1.2 '
        'for the two-dimensional case.'
    )

    doc.add_paragraph(
        'Subsequent work has identified SOC-like behavior in a diverse range of natural and '
        'artificial systems, including earthquakes (Gutenberg-Richter law), neural activity '
        '(Beggs & Plenz, 2003), forest fires (Drossel & Schwabl, 1992), and financial markets '
        '(Mantegna & Stanley, 1999). However, the universality of SOC has been questioned by '
        'several researchers who argue that many putative examples of SOC are better explained by '
        'conventional critical phenomena with fine-tuned parameters (Frigg, 2003; Watkins et al., 2016).'
    )

    doc.add_heading('2.3 Network Theory and Adaptive Systems', level=2)

    doc.add_paragraph(
        'The topology of interactions plays a crucial role in determining the emergent properties '
        'of complex systems. The discovery of small-world networks (Watts & Strogatz, 1998) and '
        'scale-free networks (Barabasi & Albert, 1999) revolutionized our understanding of how '
        'network structure influences dynamical processes such as information propagation, epidemic '
        'spreading, and synchronization (Newman, 2010).'
    )

    doc.add_paragraph(
        'Of particular relevance to this dissertation is the concept of adaptive networks, where '
        'the topology itself evolves in response to the dynamical state of the system. Gross and '
        'Blasius (2008) provided a comprehensive review of adaptive coevolutionary networks, '
        'highlighting how the interplay between dynamics and topology can give rise to novel '
        'emergent phenomena, including spontaneous formation of communities, hierarchical organization, '
        'and multistability.'
    )

    # --- Page break: Chapter 3 ---
    doc.add_page_break()

    doc.add_heading('Chapter 3: Theoretical Framework', level=1)

    doc.add_heading('3.1 The Adaptive Coupling Framework', level=2)

    doc.add_paragraph(
        'We now present the Adaptive Coupling Framework (ACF), the central theoretical '
        'contribution of this dissertation. Consider a system of N agents, each characterized by '
        'a state variable x_i(t) belonging to R^d, where d is the dimensionality of the state space. '
        'Agents interact through a weighted, directed network G = (V, E, W), where V = {1, ..., N} '
        'is the set of agents, E is the set of directed edges, and W: E -> R+ assigns coupling '
        'weights to each edge.'
    )

    doc.add_paragraph(
        'The dynamics of agent i are governed by the following stochastic differential equation:'
    )

    eq = doc.add_paragraph()
    eq.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = eq.add_run('dx_i = [f(x_i) + sum_j w_ij * g(x_i, x_j)] dt + sigma * dB_i(t)')
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'where f: R^d -> R^d represents the intrinsic dynamics of each agent, g: R^d x R^d -> R^d '
        'is the coupling function, w_ij is the adaptive coupling weight from agent j to agent i, '
        'sigma > 0 is the noise intensity, and B_i(t) is a d-dimensional Brownian motion. The key '
        'innovation of ACF lies in the evolution equation for the coupling weights:'
    )

    eq2 = doc.add_paragraph()
    eq2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = eq2.add_run('dw_ij = alpha * [phi(x_i, x_j) - w_ij] dt + eta * dW_ij(t)')
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'Here, alpha > 0 is the adaptation rate, phi: R^d x R^d -> R+ is a similarity function '
        'that measures the affinity between agents i and j, eta > 0 is the weight noise intensity, '
        'and W_ij(t) is an independent Brownian motion. This formulation captures the essential '
        'feature of adaptive systems: agents that exhibit similar behavior tend to strengthen their '
        'connections, while dissimilar agents weaken their links.'
    )

    doc.add_heading('3.2 Mean-Field Reduction', level=2)

    doc.add_paragraph(
        'To analyze the macroscopic behavior of the ACF system, we employ a generalized mean-field '
        'reduction that accounts for the heterogeneity in coupling weights. Let rho(x, w, t) denote '
        'the joint probability density of finding an agent in state x with average coupling weight w '
        'at time t. In the limit N -> infinity, the evolution of rho is governed by a '
        'Fokker-Planck equation that we derive in Appendix A.'
    )

    doc.add_paragraph(
        'The critical insight is that the mean-field reduction reveals a bifurcation structure '
        'that depends on both the average coupling strength and the variance of the coupling '
        'distribution. Specifically, we show that the order parameter psi, defined as the '
        'magnitude of the first Fourier mode of the spatial density, undergoes a continuous '
        'phase transition at a critical coupling strength that is modulated by the adaptation '
        'rate alpha and noise intensities sigma and eta.'
    )

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
