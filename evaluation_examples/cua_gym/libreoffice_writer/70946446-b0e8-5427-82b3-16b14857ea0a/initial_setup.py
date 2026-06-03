"""
Initial Setup: Create a 4-page newsletter draft with 10 paragraphs, no sections
Task ID: writer_struct_055
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_055'
OUTPUT = f'{WORKDIR}/newsletter_draft.docx'
DESKTOP_OUTPUT = f'{WORKDIR}/Desktop/newsletter_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page size to Letter (8.5 x 11 inches)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('CITY CONNECT NEWSLETTER')
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = subtitle.add_run('Monthly Community Update — March 2025')
    run2.font.size = Pt(12)
    run2.italic = True

    doc.add_paragraph()  # blank line

    # Lead Stories Header
    lead_hdr = doc.add_paragraph()
    run3 = lead_hdr.add_run('LEAD STORIES')
    run3.bold = True
    run3.font.size = Pt(14)

    doc.add_paragraph()  # blank line

    # Paragraph 1 — Lead Story 1
    p1 = doc.add_paragraph()
    run_p1 = p1.add_run(
        'Downtown Revitalization Project Breaks Ground: The long-awaited downtown revitalization '
        'project officially broke ground this week, marking a major milestone for City Connect. '
        'Mayor Linda Hargrove addressed hundreds of residents at the ceremony, emphasizing the '
        'project\'s economic potential. Construction is expected to span 18 months and will '
        'include new pedestrian walkways, green spaces, and upgraded street lighting throughout '
        'the business district.'
    )
    run_p1.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 2 — Lead Story 2
    p2 = doc.add_paragraph()
    run_p2 = p2.add_run(
        'New Public Library Branch Opens in Eastside Neighborhood: Residents of the Eastside '
        'neighborhood celebrated the opening of a brand-new public library branch on Saturday. '
        'The facility spans 12,000 square feet and features a dedicated children\'s wing, '
        'co-working spaces for remote workers, and digital resource centers equipped with the '
        'latest technology. Library Director Samuel Okafor noted that circulation figures are '
        'already exceeding initial projections during the first week of operation.'
    )
    run_p2.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 3 — Lead Story 3
    p3 = doc.add_paragraph()
    run_p3 = p3.add_run(
        'City Council Approves $4.2 Million Infrastructure Budget: At Monday night\'s city '
        'council session, members voted unanimously to approve a $4.2 million infrastructure '
        'improvement budget for the upcoming fiscal year. The funds will be allocated across '
        'road repairs, bridge maintenance, and stormwater management upgrades. Council member '
        'Patricia Weiss stated that this investment represents a commitment to long-term '
        'sustainability and resident quality of life across all neighborhoods.'
    )
    run_p3.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 4 — Lead Story 4
    p4 = doc.add_paragraph()
    run_p4 = p4.add_run(
        'Annual Community Health Fair Draws Record Attendance: This year\'s Annual Community '
        'Health Fair at Riverside Park attracted over 3,500 visitors, surpassing last year\'s '
        'attendance by nearly 40 percent. Dozens of healthcare providers offered free screenings '
        'for blood pressure, diabetes, and vision. Organizer Dr. Mei Lin Tanaka praised the '
        'community\'s enthusiasm and announced plans to expand the event to a two-day format '
        'next year to accommodate growing interest.'
    )
    run_p4.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 5 — Lead Story 5
    p5 = doc.add_paragraph()
    run_p5 = p5.add_run(
        'New Transit Line Proposal Advances to Public Comment Phase: The Metropolitan Transit '
        'Authority has advanced its proposal for a new light rail line connecting the airport '
        'to the city center into the public comment phase. The proposed route would span '
        '14.3 miles and include eight stations, with an estimated ridership of 28,000 daily '
        'commuters. Public comment sessions are scheduled throughout April, and residents are '
        'encouraged to attend and share their feedback on route planning and station locations.'
    )
    run_p5.font.size = Pt(11)

    doc.add_paragraph()  # blank line before page break

    # Page break before brief items section
    doc.add_page_break()

    # Brief Items Header
    brief_hdr = doc.add_paragraph()
    run_bh = brief_hdr.add_run('BRIEF ITEMS')
    run_bh.bold = True
    run_bh.font.size = Pt(14)

    doc.add_paragraph()  # blank line

    # Paragraph 6 — Brief Item 1
    p6 = doc.add_paragraph()
    run_p6 = p6.add_run(
        'Community Garden Volunteer Day: The Maple Street Community Garden is hosting a spring '
        'clean-up and planting day on April 5th from 9 AM to 1 PM. Volunteers are welcome; '
        'tools and refreshments will be provided. Contact coordinator Rosa Fuentes to sign up.'
    )
    run_p6.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 7 — Brief Item 2
    p7 = doc.add_paragraph()
    run_p7 = p7.add_run(
        'Youth Soccer League Registration Open: Registration for the spring youth soccer league '
        'is now open for children ages 6 to 14. The season runs from May through July, with '
        'games held at Greenfield Sports Complex. Visit cityconnect.org/soccer to register '
        'before the April 15th deadline.'
    )
    run_p7.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 8 — Brief Item 3
    p8 = doc.add_paragraph()
    run_p8 = p8.add_run(
        'Noise Ordinance Reminder: With warmer weather approaching, the city reminds residents '
        'that the noise ordinance prohibits amplified sound after 10 PM on weekdays and '
        '11 PM on weekends. Repeat violations may result in fines starting at $150. '
        'Questions can be directed to the City Compliance Office at ext. 4412.'
    )
    run_p8.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 9 — Brief Item 4
    p9 = doc.add_paragraph()
    run_p9 = p9.add_run(
        'Fire Station 7 Open House: Fire Station 7 on Birch Avenue will host an open house '
        'on April 12th from 10 AM to 2 PM. Families are invited to tour the station, meet '
        'firefighters, and learn about fire safety. Children under 12 can try on gear and '
        'take photos aboard Engine 7.'
    )
    run_p9.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # Paragraph 10 — Brief Item 5
    p10 = doc.add_paragraph()
    run_p10 = p10.add_run(
        'Pothole Reporting App Launched: The city has launched a new mobile app allowing '
        'residents to report potholes and road damage directly from their smartphones. '
        'Reports are geotagged and routed automatically to the Public Works Department. '
        'The app, "Fix My Street," is available for free on iOS and Android platforms.'
    )
    run_p10.font.size = Pt(11)

    # Ensure the Desktop directory exists and save there
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc.save(DESKTOP_OUTPUT)
    print(f'Initial file created: {DESKTOP_OUTPUT}')

    # GUI-ready startup
    # Kill any existing LibreOffice to avoid stale locks
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(2)

    launch_gui(f'libreoffice --writer "{DESKTOP_OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
