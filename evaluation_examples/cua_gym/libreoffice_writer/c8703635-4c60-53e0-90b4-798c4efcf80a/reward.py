"""
Reward Script: Remove bullet formatting from all items and convert to plain paragraphs
Task ID: writer_list_006
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All 5 target paragraphs have style changed away from 'List Bullet'
  Component 2 (0.3 pts): No list-related XML (numPr or ListBullet pStyle) remains in target paragraphs
  Component 3 (0.2 pts): All 5 paragraph texts preserved intact AND paragraphs are non-list styled
                          (compound check - ensures text and conversion are both present)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_list_006'

# The five expected paragraph texts (ground truth from task context)
EXPECTED_TEXTS = [
    'Consider restructuring the team',
    'Update the project timeline by Friday',
    'Send progress report to stakeholders',
    'Book conference room for demo',
    'Prepare slide deck for presentation',
]


def verify_task(file_path):
    """
    Verify that bullet list formatting has been removed from all five target paragraphs.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify document has enough paragraphs
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected at least 7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Identify the target paragraphs by matching expected texts
    target_paras = []
    for para in doc.paragraphs:
        if para.text.strip() in EXPECTED_TEXTS:
            target_paras.append(para)

    print(f"Found {len(target_paras)} target paragraphs out of expected 5")

    if len(target_paras) == 0:
        print("FAIL: No target paragraphs found — text content may have been lost")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 5 target paragraphs have style changed away from 'List Bullet' (0.5 points)
    # In initial_env all 5 have style='List Bullet'; in golden_env they should be 'Normal'
    # FAILS on initial_env (style='List Bullet') → PASSES on golden_env (style='Normal')
    try:
        not_list_bullet_count = 0
        for para in target_paras:
            style_name = para.style.name
            # A list bullet style has these typical names in python-docx
            is_list_style = 'List' in style_name or 'Bullet' in style_name
            if not is_list_style:
                not_list_bullet_count += 1
                print(f"PASS: Paragraph '{para.text[:40]}' style={style_name!r} (not list)")
            else:
                print(f"FAIL: Paragraph '{para.text[:40]}' still has list-type style={style_name!r}")

        if not_list_bullet_count == len(target_paras) and len(target_paras) == 5:
            print(f"PASS: Component 1 — All 5 paragraphs have non-list styles (0.5 pts)")
            total_score += 0.5
        elif not_list_bullet_count > 0:
            partial = round((not_list_bullet_count / 5) * 0.5, 2)
            print(f"PARTIAL: Component 1 — {not_list_bullet_count}/5 paragraphs converted ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — 0/5 paragraphs converted away from List Bullet style")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: No list-related XML (numPr or ListBullet pStyle) in target paragraphs (0.3 points)
    # In initial_env: pPr contains <w:pStyle w:val="ListBullet"/>
    # In golden_env: no pPr or pPr has no list-related elements
    # FAILS on initial_env → PASSES on golden_env
    try:
        no_list_xml_count = 0
        for para in target_paras:
            pPr = para._element.pPr

            # Count list-related XML elements: numPr elements + list-type pStyle elements
            numPr_count = 0
            list_pstyle_count = 0
            if pPr is not None:
                # Check for numPr (explicit numbering reference)
                numPr_elem = pPr.find(qn('w:numPr'))
                if numPr_elem is not None:
                    numPr_count += 1
                    print(f"FAIL: Paragraph '{para.text[:40]}' still has numPr in pPr")

                # Check for pStyle pointing to a list style
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'), '')
                    if 'List' in style_val or 'Bullet' in style_val or 'Number' in style_val:
                        list_pstyle_count += 1
                        print(f"FAIL: Paragraph '{para.text[:40]}' pStyle={style_val!r} (list-related)")

            if numPr_count == 0 and list_pstyle_count == 0:
                no_list_xml_count += 1
                print(f"PASS: Paragraph '{para.text[:40]}' has no list XML in pPr")

        if no_list_xml_count == len(target_paras) and len(target_paras) == 5:
            print(f"PASS: Component 2 — All 5 paragraphs have no list-related XML (0.3 pts)")
            total_score += 0.3
        elif no_list_xml_count > 0:
            partial = round((no_list_xml_count / 5) * 0.3, 2)
            print(f"PARTIAL: Component 2 — {no_list_xml_count}/5 paragraphs free of list XML ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — All paragraphs still have list-related XML")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All 5 paragraph texts preserved AND each converted paragraph has non-list style (0.2 pts)
    # Compound check: text exists + style is non-list → only BOTH conditions together earn points
    # FAILS on initial_env (style='List Bullet') → PASSES on golden_env (style='Normal' + text intact)
    try:
        doc_texts_with_styles = {
            para.text.strip(): para.style.name
            for para in doc.paragraphs
        }
        preserved_and_converted = 0
        for expected_text in EXPECTED_TEXTS:
            if expected_text in doc_texts_with_styles:
                style_name = doc_texts_with_styles[expected_text]
                is_list_style = 'List' in style_name or 'Bullet' in style_name
                if not is_list_style:
                    preserved_and_converted += 1
                    print(f"PASS: Text '{expected_text[:40]}' preserved with non-list style={style_name!r}")
                else:
                    print(f"FAIL: Text '{expected_text[:40]}' present but still has list style={style_name!r}")
            else:
                print(f"FAIL: Text '{expected_text[:40]}' not found in document")

        if preserved_and_converted == 5:
            print(f"PASS: Component 3 — All 5 texts preserved with non-list styles (0.2 pts)")
            total_score += 0.2
        elif preserved_and_converted > 0:
            partial = round((preserved_and_converted / 5) * 0.2, 2)
            print(f"PARTIAL: Component 3 — {preserved_and_converted}/5 texts preserved and converted ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No paragraphs pass compound text+style check")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/notes_draft.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
