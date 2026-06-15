"""
Initial Setup: Remove bullet formatting from all items in document
Task ID: writer_list_006
Domain: libreoffice_writer

Creates a .docx file with 5 paragraphs formatted as a bulleted list.
The agent task is to remove the bullet formatting and convert them to plain paragraphs.
"""

import os
import shlex
import subprocess
import time
from docx import Document

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_list_006'
OUTPUT = f'{WORKDIR}/notes_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title/heading for context
    doc.add_heading("Meeting Notes", level=1)

    # Add a brief introductory paragraph
    doc.add_paragraph("The following action items were identified during today's planning session:")

    # Add 5 paragraphs formatted as bulleted list (List Bullet style)
    bullet_items = [
        "Consider restructuring the team",
        "Update the project timeline by Friday",
        "Send progress report to stakeholders",
        "Book conference room for demo",
        "Prepare slide deck for presentation",
    ]

    for item in bullet_items:
        doc.add_paragraph(item, style="List Bullet")

    # Add a closing paragraph after the list
    doc.add_paragraph("Please follow up on all action items before the next meeting.")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
