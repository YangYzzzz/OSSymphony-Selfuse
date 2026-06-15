"""
Initial Setup: Research Findings document with track changes off, no footnotes.
Task ID: writer_rm_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Research Findings Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Quarterly Consumer Behavior Study — Q1-Q2 2025')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Consumer behavior patterns have shifted significantly over the past decade, '
        'driven by the rise of digital commerce and evolving demographic preferences. '
        'This report presents the findings of a comprehensive study conducted by the '
        'Market Research Division at Clearwater Analytics during the first half of 2025.'
    )
    doc.add_paragraph(
        'The study was commissioned to evaluate purchasing trends across four major '
        'retail sectors: electronics, apparel, home goods, and grocery. Our research '
        'team, led by Dr. Elena Marchetti and Senior Analyst Raj Patel, designed the '
        'methodology to capture both quantitative metrics and qualitative sentiment data.'
    )

    # --- Section 2: Methodology ---
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        'The research employed a mixed-methods approach combining structured online '
        'surveys with follow-up telephone interviews. Survey participants were recruited '
        'from a nationally representative panel maintained by DataPulse Research Inc., '
        'stratified by age, income, and geographic region.'
    )
    doc.add_paragraph(
        'Quantitative data were analyzed using multivariate regression models to '
        'identify statistically significant predictors of purchasing frequency. '
        'Qualitative responses were coded by two independent reviewers using thematic '
        'analysis, with an inter-rater reliability coefficient of 0.87.'
    )

    # --- Section 3: Data Analysis ---
    doc.add_heading('3. Data Analysis', level=1)
    doc.add_paragraph(
        'The primary dataset comprised 2,500 survey responses collected between '
        'January and June 2025. After data cleaning and validation, 2,347 responses '
        'met the inclusion criteria and were retained for analysis.'
    )
    doc.add_paragraph(
        'Demographic breakdowns revealed that 54% of respondents were female, 43% '
        'male, and 3% identified as non-binary. The median age was 38 years, with '
        'a standard deviation of 12.4 years. Household income ranged from $24,000 to '
        '$185,000, with a median of $67,500.'
    )
    doc.add_paragraph(
        'Cross-tabulation analysis showed strong correlations between income brackets '
        'and electronics spending. Respondents in the top quartile (income above $112,000) '
        'spent an average of $3,240 annually on electronics, compared to $890 for those '
        'in the bottom quartile.'
    )

    # --- Section 4: Results ---
    doc.add_heading('4. Results', level=1)
    doc.add_paragraph(
        'Key findings indicate a 17.3% increase in online grocery spending among '
        'respondents aged 25-44, compared to the same period in 2024. This shift '
        'was particularly pronounced in urban areas with populations exceeding 500,000, '
        'where convenience and delivery speed were cited as primary motivators by 68% '
        'of respondents.'
    )
    doc.add_paragraph(
        'The apparel sector exhibited the most volatility, with seasonal spending '
        'fluctuations of up to 42% between Q1 and Q2. Brand loyalty metrics declined '
        'by 8.6 percentage points year-over-year, suggesting increased price sensitivity '
        'among consumers. Notably, 61% of respondents reported using at least three '
        'different comparison platforms before making apparel purchases over $75.'
    )

    # --- Section 5: Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        'The findings of this study underscore the accelerating transformation of '
        'consumer purchasing behavior in the post-pandemic retail landscape. The data '
        'support the hypothesis that digital-first strategies are no longer optional '
        'for retailers targeting the 25-54 age demographic. We recommend that stakeholders '
        'review the detailed appendices for sector-specific strategic recommendations '
        'and projected trend models through 2026.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
