"""
Initial Setup: Create a sale flyer document without decorative header.
Task ID: writer_frd_082
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_082'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # -- No decorative header / Fontwork -- the task asks the agent to add one --

    # Spacer at the top (where the Fontwork will go)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(24)

    # Subheading
    sub = doc.add_paragraph()
    sub.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.paragraph_format.space_after = Pt(6)
    run_sub = sub.add_run("Up to 60% Off on Selected Items!")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(0xE6, 0x4A, 0x19)
    run_sub.italic = True

    # Date line
    date_line = doc.add_paragraph()
    date_line.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_line.paragraph_format.space_after = Pt(18)
    run_date = date_line.add_run("June 15 - July 31, 2025  |  All Locations")
    run_date.font.size = Pt(11)
    run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Horizontal rule (simulated with underscores)
    hr = doc.add_paragraph()
    hr.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hr.paragraph_format.space_after = Pt(12)
    run_hr = hr.add_run("_" * 60)
    run_hr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run_hr.font.size = Pt(8)

    # Section: Electronics
    h_elec = doc.add_heading("Electronics & Gadgets", level=2)
    for run in h_elec.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    products_electronics = [
        ("Meridian 65\" 4K Smart TV", "$899.99", "$1,499.99", "40% Off"),
        ("AuraSound Noise-Cancelling Headphones", "$129.95", "$249.99", "48% Off"),
        ("SwiftCharge 20000mAh Power Bank", "$34.99", "$59.99", "42% Off"),
        ("ProView 27\" Gaming Monitor 165Hz", "$329.00", "$549.00", "40% Off"),
    ]

    table_elec = doc.add_table(rows=1, cols=4)
    table_elec.style = "Table Grid"
    hdr_cells = table_elec.rows[0].cells
    headers = ["Product", "Sale Price", "Original Price", "Discount"]
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        shading = hdr_cells[i]._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1A5C8A',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elem)

    for prod in products_electronics:
        row = table_elec.add_row()
        for i, val in enumerate(prod):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 1:
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            elif i == 2:
                run.font.strike = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacer

    # Section: Home & Living
    h_home = doc.add_heading("Home & Living", level=2)
    for run in h_home.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    products_home = [
        ("CloudRest Memory Foam Mattress (Queen)", "$449.00", "$799.00", "44% Off"),
        ("BreezeMax Tower Fan with Remote", "$64.99", "$119.99", "46% Off"),
        ("ArtisanBrew Coffee Maker 12-Cup", "$49.95", "$89.95", "44% Off"),
        ("LuxeThread 800TC Egyptian Cotton Sheet Set", "$79.99", "$159.99", "50% Off"),
        ("ScentWave Ultrasonic Diffuser & Oils Kit", "$27.50", "$44.99", "39% Off"),
    ]

    table_home = doc.add_table(rows=1, cols=4)
    table_home.style = "Table Grid"
    hdr_cells2 = table_home.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr_cells2[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = hdr_cells2[i]._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1A5C8A',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elem)

    for prod in products_home:
        row = table_home.add_row()
        for i, val in enumerate(prod):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 1:
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            elif i == 2:
                run.font.strike = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacer

    # Section: Outdoor & Sports
    h_outdoor = doc.add_heading("Outdoor & Sports", level=2)
    for run in h_outdoor.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    products_outdoor = [
        ("TrailBlaze 50L Hiking Backpack", "$59.99", "$99.99", "40% Off"),
        ("AquaGlide Inflatable Kayak (2-Person)", "$189.00", "$349.00", "46% Off"),
        ("SunShield UV50+ Beach Tent", "$39.95", "$69.95", "43% Off"),
    ]

    table_outdoor = doc.add_table(rows=1, cols=4)
    table_outdoor.style = "Table Grid"
    hdr_cells3 = table_outdoor.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr_cells3[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = hdr_cells3[i]._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1A5C8A',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elem)

    for prod in products_outdoor:
        row = table_outdoor.add_row()
        for i, val in enumerate(prod):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 1:
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            elif i == 2:
                run.font.strike = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacer

    # Footer info
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para.paragraph_format.space_before = Pt(18)
    run_f1 = footer_para.add_run("Visit us at: ")
    run_f1.font.size = Pt(10)
    run_f2 = footer_para.add_run("www.summersalecentral.com")
    run_f2.font.size = Pt(10)
    run_f2.bold = True
    run_f2.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    terms = doc.add_paragraph()
    terms.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_t = terms.add_run(
        "Prices valid while supplies last. Cannot be combined with other offers. "
        "See store for details."
    )
    run_t.font.size = Pt(8)
    run_t.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_t.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
