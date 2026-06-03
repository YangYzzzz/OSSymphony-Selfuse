"""
Reward Script: Insert Fontwork (WordArt) 'SUMMER SALE' with curved wave style at top of flyer
Task ID: writer_frd_082
Domain: libreoffice_writer
Scoring:
  Component 1: WordArt/Fontwork drawing element exists (0.25)
  Component 2: Text reads 'SUMMER SALE' (0.30)
  Component 3: Wave/curve text warp style applied (0.25)
  Component 4: Positioned at the top of the document (0.20)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_082'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Namespace map for XML parsing
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    }

    body = doc.element.body

    # ---- Gather all drawing elements with WordArt characteristics ----
    wordart_elements = []
    wordart_para_indices = []

    for i, para in enumerate(doc.paragraphs):
        p_elem = para._element
        # Check both inline and anchor drawings
        for drawing_type in ('inline', 'anchor'):
            for drawing in p_elem.iter():
                tag = drawing.tag.split('}')[-1] if '}' in drawing.tag else drawing.tag
                if tag != drawing_type:
                    continue
                # Check if this is a WordArt (fromWordArt="1" in bodyPr, or docPr name contains WordArt/Fontwork)
                is_wordart = False

                # Method 1: Check bodyPr fromWordArt attribute
                for bp in drawing.iter():
                    bp_tag = bp.tag.split('}')[-1] if '}' in bp.tag else bp.tag
                    if bp_tag == 'bodyPr' and bp.attrib.get('fromWordArt') == '1':
                        is_wordart = True
                        break

                # Method 2: Check docPr name
                if not is_wordart:
                    for dp in drawing.iter():
                        dp_tag = dp.tag.split('}')[-1] if '}' in dp.tag else dp.tag
                        if dp_tag == 'docPr':
                            name = dp.attrib.get('name', '').lower()
                            if 'wordart' in name or 'fontwork' in name:
                                is_wordart = True
                                break

                # Method 3: Check for prstTxWarp (text warp = Fontwork even without explicit flag)
                if not is_wordart:
                    for tw in drawing.iter():
                        tw_tag = tw.tag.split('}')[-1] if '}' in tw.tag else tw.tag
                        if tw_tag == 'prstTxWarp':
                            prst = tw.attrib.get('prst', '')
                            if prst and prst != 'textNoShape':
                                is_wordart = True
                                break

                if is_wordart:
                    wordart_elements.append(drawing)
                    wordart_para_indices.append(i)

    # Component 1: WordArt/Fontwork drawing element exists (0.25 points)
    try:
        if len(wordart_elements) > 0:
            print(f"PASS: Component 1 — WordArt/Fontwork element found ({len(wordart_elements)} element(s)) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — No WordArt/Fontwork drawing element found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    if len(wordart_elements) == 0:
        # No WordArt found, all subsequent checks will fail
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Use the first WordArt element for remaining checks
    wordart = wordart_elements[0]
    wordart_para_idx = wordart_para_indices[0]

    # Component 2: Text reads 'SUMMER SALE' (0.30 points)
    try:
        # Extract text from the WordArt element
        found_texts = []
        for t_elem in wordart.iter():
            t_tag = t_elem.tag.split('}')[-1] if '}' in t_elem.tag else t_elem.tag
            if t_tag == 't' and t_elem.text:
                found_texts.append(t_elem.text.strip())

        # Deduplicate (XML may repeat text in different contexts)
        unique_texts = list(set(found_texts))
        combined_text = ' '.join(unique_texts).strip().upper()

        # Check if SUMMER SALE is present
        has_summer_sale = False
        for txt in unique_texts:
            if txt.strip().upper() == 'SUMMER SALE':
                has_summer_sale = True
                break
        if not has_summer_sale:
            # Also check combined
            if 'SUMMER SALE' in combined_text:
                has_summer_sale = True

        if has_summer_sale:
            print(f"PASS: Component 2 — WordArt text is 'SUMMER SALE' (found: {unique_texts}) (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — Expected text 'SUMMER SALE', found: {unique_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Wave/curve text warp style applied (0.25 points)
    try:
        warp_preset = None
        for tw in wordart.iter():
            tw_tag = tw.tag.split('}')[-1] if '}' in tw.tag else tw.tag
            if tw_tag == 'prstTxWarp':
                warp_preset = tw.attrib.get('prst', '')
                break

        # Accept any wave-type warp: textWave1, textWave2, textWave4, textDeflate, textInflate, etc.
        wave_presets = {
            'textWave1', 'textWave2', 'textWave4',
            'textCurveUp', 'textCurveDown',
            'textDeflate', 'textInflate',
            'textDoubleWave1',
        }
        # Also accept any preset that contains 'wave' or 'curve' (case-insensitive)
        is_wave = False
        if warp_preset:
            if warp_preset in wave_presets:
                is_wave = True
            elif 'wave' in warp_preset.lower() or 'curve' in warp_preset.lower():
                is_wave = True

        if is_wave:
            print(f"PASS: Component 3 — Wave/curve warp style applied (preset: {warp_preset}) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Expected wave/curve warp style, found preset: {warp_preset}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Positioned at the top of the document (0.20 points)
    try:
        num_paras = len(doc.paragraphs)
        # The WordArt should be in one of the first few paragraphs (top of document)
        # Allow up to paragraph index 2 (first 3 paragraphs) for "top" positioning
        if wordart_para_idx <= 2:
            print(f"PASS: Component 4 — WordArt is at top of document (paragraph {wordart_para_idx}) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — WordArt is at paragraph {wordart_para_idx} (expected within first 3 paragraphs for 'top' positioning)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
