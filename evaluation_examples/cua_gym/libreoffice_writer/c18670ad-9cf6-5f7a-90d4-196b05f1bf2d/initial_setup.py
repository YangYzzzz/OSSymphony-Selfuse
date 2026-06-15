"""
Initial Setup: Parish Newsletter with 2-column layout (spacing=0.50cm, no separator)
Task ID: writer_page_032
Domain: libreoffice_writer

Creates /home/user/Desktop/parish_newsletter.docx with:
- 4-page A4 parish newsletter
- Page margins: top=2.0cm, bottom=2.0cm, left=1.5cm, right=1.5cm
- Two equal-width columns, spacing=0.50cm, NO separator line
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
FILENAME = 'parish_newsletter.docx'
OUTPUT = f'{WORKDIR}/{FILENAME}'

# Conversion: 1 cm = 567 twips (OOXML column widths use twips)
CM_TO_TWIPS = 567


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_two_columns(section, spacing_cm=0.50, sep=False):
    """
    Apply a 2-column layout to a section via raw XML.
    spacing_cm: gap between the two columns in cm
    sep: whether to add a vertical separator line between columns
    """
    sect_pr = section._sectPr

    # Remove any existing cols element
    existing = sect_pr.find(qn('w:cols'))
    if existing is not None:
        sect_pr.remove(existing)

    # Page printable width: A4=21cm, left=1.5cm, right=1.5cm => 18cm
    printable_cm = 18.0
    space_twips = int(spacing_cm * CM_TO_TWIPS)
    col_width_cm = (printable_cm - spacing_cm) / 2.0
    col_width_twips = int(col_width_cm * CM_TO_TWIPS)

    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), str(space_twips))
    if sep:
        cols.set(qn('w:sep'), '1')

    col1 = OxmlElement('w:col')
    col1.set(qn('w:w'), str(col_width_twips))
    col1.set(qn('w:space'), str(space_twips))

    col2 = OxmlElement('w:col')
    col2.set(qn('w:w'), str(col_width_twips))

    cols.append(col1)
    cols.append(col2)

    # Insert cols before pgMar or at end of sectPr
    pg_mar = sect_pr.find(qn('w:pgMar'))
    if pg_mar is not None:
        pg_mar.addprevious(cols)
    else:
        sect_pr.append(cols)


def add_heading(doc, text, level=1, color=None):
    """Add a heading paragraph."""
    para = doc.add_heading(text, level=level)
    if color:
        for run in para.runs:
            run.font.color.rgb = color
    return para


def add_body_text(doc, text, bold=False, italic=False):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return para


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # ---- Page Setup ----
    section = doc.sections[0]
    section.page_width = Cm(21)        # A4
    section.page_height = Cm(29.7)     # A4 portrait
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ---- Two-column layout: spacing=0.50cm, no separator line ----
    set_two_columns(section, spacing_cm=0.50, sep=False)

    # ---- Content: 4-page Parish Newsletter ----

    # --- PAGE 1: Masthead and Welcome ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("ST. MARGARET'S PARISH NEWSLETTER")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run("Volume 24 · Issue 3 · March 2025")
    date_run.italic = True
    date_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    add_heading(doc, "A Message from Father Collins", level=2)

    add_body_text(doc,
        "Dear parishioners, as we enter the season of Lent, I invite each of you to "
        "reflect deeply on the blessings God has bestowed upon our community. This "
        "year's Lenten journey calls us to prayer, fasting, and almsgiving — three "
        "pillars that unite us in faith and purpose."
    )
    add_body_text(doc,
        "Our parish has grown remarkably over the past year. We welcomed 47 new "
        "families to St. Margaret's, celebrated 12 baptisms, and saw our youth group "
        "swell to over 80 active young members. These are signs of a living, "
        "breathing community of faith."
    )
    add_body_text(doc,
        "I am particularly grateful to the many volunteers who give so freely of their "
        "time: the choir under the direction of Mrs. Patricia Hennessy, the Legion of "
        "Mary chapter led by Mr. Thomas Rafferty, and our dedicated sacristans who "
        "ensure every liturgy proceeds with reverence and dignity."
    )

    add_heading(doc, "Lenten Schedule 2025", level=2)

    add_body_text(doc, "Wednesday Masses: 7:00 AM and 7:30 PM")
    add_body_text(doc, "Stations of the Cross: Every Friday at 7:00 PM")
    add_body_text(doc, "Confessions: Saturday 3:00–5:00 PM and by appointment")
    add_body_text(doc, "Ash Wednesday (5 March): Masses at 8:00 AM, 12:10 PM, and 7:30 PM")
    add_body_text(doc, "Palm Sunday (13 April): Procession begins at 9:45 AM")
    add_body_text(doc, "Holy Thursday (17 April): Mass of the Lord's Supper at 7:30 PM")
    add_body_text(doc, "Good Friday (18 April): Liturgy at 3:00 PM")
    add_body_text(doc, "Easter Vigil (19 April): Begins at 8:30 PM")
    add_body_text(doc, "Easter Sunday (20 April): Masses at 8:00 AM, 10:00 AM, 12:00 PM")

    doc.add_page_break()

    # --- PAGE 2: Community News and Events ---
    add_heading(doc, "Community News", level=1)

    add_heading(doc, "Annual Parish Fundraising Dinner", level=2)
    add_body_text(doc,
        "The 18th Annual Parish Fundraising Dinner will be held on Saturday, "
        "29 March 2025, at 6:30 PM in St. Margaret's Hall. Tickets are €35 per "
        "person and include a three-course meal with live music provided by the "
        "Clancy Street Band. All proceeds go towards the renovation of our parish "
        "centre roof, which has required extensive repairs following last winter's "
        "storms."
    )
    add_body_text(doc,
        "Tickets are available after Sunday Masses or from the parish office. "
        "Please contact Mary O'Brien at 087-555-0123 or email "
        "events@stmargarets.ie. Tables of eight can be reserved in advance. "
        "We hope to see you there for a wonderful evening of food, music, and "
        "fellowship."
    )

    add_heading(doc, "New Ministries Being Launched", level=2)
    add_body_text(doc,
        "We are thrilled to announce the launch of two new parish ministries this "
        "spring. The first is a Bereavement Support Group, which will meet on the "
        "second Monday of each month at 7:30 PM in the parish meeting room. The "
        "group is facilitated by trained volunteers and offers a safe, welcoming "
        "space for those experiencing grief."
    )
    add_body_text(doc,
        "The second new ministry is a Parish Garden Project, which aims to create "
        "a community vegetable and herb garden on the grounds behind the church "
        "hall. Produce from the garden will be donated to the local St. Vincent de "
        "Paul Society food bank. Volunteers of all ages are welcome — no gardening "
        "experience required!"
    )

    add_heading(doc, "Youth Group Updates", level=2)
    add_body_text(doc,
        "Our youth group, led by Coordinator Aoife Brennan, has had a remarkable "
        "few months. In February, 24 young people from our parish took part in a "
        "regional youth pilgrimage to Knock, County Mayo. Participants reported it "
        "as a deeply moving experience and many are already looking forward to "
        "next year's trip."
    )
    add_body_text(doc,
        "The youth group is also preparing for the Diocesan Drama Festival, where "
        "they will present an original production titled 'The Journey of Simeon' "
        "written by group member Ciarán Murphy (age 17). Rehearsals are every "
        "Tuesday evening from 7:00–9:00 PM. Supporters are welcome to attend the "
        "festival performance on 26 April."
    )

    doc.add_page_break()

    # --- PAGE 3: Parish Services and Notices ---
    add_heading(doc, "Parish Services", level=1)

    add_heading(doc, "Mass Times", level=2)
    add_body_text(doc, "Sunday: 8:00 AM, 10:00 AM (Family Mass), 12:00 PM, 6:30 PM")
    add_body_text(doc, "Monday to Friday: 8:00 AM and 10:00 AM")
    add_body_text(doc, "Saturday: 10:00 AM and 6:00 PM (Vigil)")
    add_body_text(doc, "Holy Days of Obligation: 8:00 AM, 10:00 AM, 12:10 PM, 7:30 PM")

    add_heading(doc, "Sacramental Preparation", level=2)
    add_body_text(doc,
        "First Holy Communion preparation classes are held every Saturday morning "
        "from 10:30 AM to 12:00 PM in Rooms 1 and 2 of the parish centre. Parents "
        "are asked to attend the first class of each month with their children. "
        "For information, contact the Parish Office."
    )
    add_body_text(doc,
        "Confirmation preparation continues with monthly evenings of prayer and "
        "reflection. The next gathering is on Tuesday, 18 March, at 7:30 PM. "
        "Candidates are reminded to bring their sponsor and faith journal."
    )

    add_heading(doc, "Parish Office Hours", level=2)
    add_body_text(doc, "Monday–Thursday: 9:00 AM – 1:00 PM and 2:00 PM – 4:30 PM")
    add_body_text(doc, "Friday: 9:00 AM – 1:00 PM")
    add_body_text(doc, "Saturday & Sunday: Closed (emergency contact available)")
    add_body_text(doc, "Phone: 01-555-0198   Email: office@stmargarets.ie")

    add_heading(doc, "Notices", level=2)
    add_body_text(doc,
        "The Parish Library will be closed for stocktaking from 10–14 March. "
        "Books due during this period may be returned in the blue drop-box at "
        "the side entrance."
    )
    add_body_text(doc,
        "A collection for the Irish Bishops' Lenten Appeal will be taken up at "
        "all Masses on the weekend of 22–23 March. Envelopes are available at "
        "the church doors."
    )
    add_body_text(doc,
        "The annual Church Gate Collection in support of the parish primary school "
        "building fund will take place on the weekend of 5–6 April. Volunteers "
        "are needed — please sign up in the church porch."
    )

    doc.add_page_break()

    # --- PAGE 4: Letters, Obituaries, and Prayer Intentions ---
    add_heading(doc, "Letters to the Parish", level=1)

    add_heading(doc, "Letter of Appreciation", level=2)
    add_body_text(doc,
        "Dear Father Collins and the parishioners of St. Margaret's, I write to "
        "express my heartfelt gratitude for the outpouring of support following "
        "the recent passing of my husband, Desmond. The kindness shown by so many "
        "— the meals brought to our door, the Mass cards, the phone calls and "
        "visits — has been a tremendous source of comfort to myself and our "
        "children during this difficult time."
    )
    add_body_text(doc,
        "Desmond was a proud member of this parish community for over 40 years. "
        "He served on the Finance Committee, volunteered at the St. Vincent de Paul "
        "collection every Christmas, and never missed a Sunday Mass unless illness "
        "prevented him. I know he would be deeply moved by the love shown by all. "
        "God bless you all.",
        italic=True
    )
    para_sig = doc.add_paragraph()
    run_sig = para_sig.add_run("— Brigid Connolly, Elm Drive")
    run_sig.italic = True
    run_sig.font.size = Pt(10)

    add_heading(doc, "Recent Bereavements", level=2)
    add_body_text(doc,
        "We remember in our prayers those who have recently passed from our "
        "parish community:"
    )
    for name in [
        "Desmond Connolly (74), Elm Drive",
        "Margaret Rose Fahey (89), The Crescent",
        "Patrick Joseph Daly (66), Meadow Lane",
        "Eileen Nora Sullivan (81), Abbeyfield Court",
    ]:
        para = doc.add_paragraph(name, style='List Bullet')
        para.runs[0].font.size = Pt(11)

    add_body_text(doc,
        "May their souls, and the souls of all the faithful departed, through the "
        "mercy of God, rest in peace. Amen."
    )

    add_heading(doc, "Prayer Intentions", level=2)
    add_body_text(doc,
        "The following intentions have been received for the coming weeks. "
        "Please hold these people and their families in your prayers:"
    )
    intentions = [
        "For the health and recovery of Seamus Byrne, who is receiving treatment in hospital.",
        "For the safe delivery of baby expected by Siobhán and Conor McCarthy.",
        "In thanksgiving for 50 years of marriage celebrated by Frank and Nora Gallagher.",
        "For students preparing for the Leaving Certificate examinations in June.",
        "For all healthcare workers serving in our local hospitals and clinics.",
        "For the success of our parish fundraising efforts this Lenten season.",
    ]
    for intention in intentions:
        para = doc.add_paragraph(intention, style='List Bullet')
        para.runs[0].font.size = Pt(10)

    # ---- Save ----
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # ---- GUI-ready startup ----
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
