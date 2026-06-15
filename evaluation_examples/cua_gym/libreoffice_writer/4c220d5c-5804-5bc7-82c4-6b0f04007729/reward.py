"""
Reward Script: Format merge field 'Salary' as currency with $ prefix, comma separators, 2 decimals
Task ID: writer_mt_018
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Salary in letter body is formatted as currency (no longer raw placeholder)
  Component 2 (0.5): All salary values in data table formatted as $X,XXX.00
  Component 3 (0.2): Letter body salary matches first candidate's formatted salary from table
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_018'

# Currency pattern: $DD,DDD.DD  (dollar sign, digits with commas, exactly 2 decimals)
CURRENCY_RE = re.compile(r'^\$[\d]{1,3}(,\d{3})*\.\d{2}$')

# Known raw salary values from the initial data source
RAW_SALARIES = ['45000', '72500', '98000', '67500', '85000']
# Expected formatted values
FORMATTED_SALARIES = ['$45,000.00', '$72,500.00', '$98,000.00', '$67,500.00', '$85,000.00']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Salary in letter body (Para 14) is formatted as currency (0.3 points)
    # Initial has "<Salary>", golden should have "$45,000.00"
    try:
        salary_para = None
        for para in doc.paragraphs:
            if 'Your annual salary will be' in para.text:
                salary_para = para.text
                break

        if salary_para is None:
            print("FAIL: Component 1 — Could not find salary paragraph")
        else:
            # Check that the paragraph does NOT contain the raw placeholder "<Salary>"
            # AND DOES contain a currency-formatted value
            has_placeholder = '<Salary>' in salary_para
            currency_match = re.search(r'\$[\d]{1,3}(,\d{3})*\.\d{2}', salary_para)

            if not has_placeholder and currency_match:
                print(f"PASS: Component 1 — Salary in body is formatted as currency: {currency_match.group()} (0.3 pts)")
                total_score += 0.3
            elif has_placeholder:
                print(f"FAIL: Component 1 — Salary paragraph still has raw placeholder '<Salary>'")
            else:
                print(f"FAIL: Component 1 — Salary paragraph has no currency-formatted value, text: {repr(salary_para)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All salary values in data table are formatted as $X,XXX.00 (0.5 points)
    # Initial table has raw numbers (45000, 72500, etc.), golden should have $45,000.00 etc.
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No tables found in document")
        else:
            table = doc.tables[0]
            # Find the Salary column index
            header_row = [cell.text.strip() for cell in table.rows[0].cells]
            if 'Salary' not in header_row:
                print(f"FAIL: Component 2 — No 'Salary' column found. Headers: {header_row}")
            else:
                sal_col = header_row.index('Salary')
                data_rows = list(table.rows)[1:]  # skip header
                total_data = len(data_rows)
                formatted_count = 0

                for ri, row in enumerate(data_rows):
                    cell_val = row.cells[sal_col].text.strip()
                    if CURRENCY_RE.match(cell_val):
                        formatted_count += 1
                    else:
                        print(f"  Row {ri+1}: Salary '{cell_val}' is NOT currency-formatted")

                if formatted_count == total_data and total_data > 0:
                    print(f"PASS: Component 2 — All {total_data} salary values in table are currency-formatted (0.5 pts)")
                    total_score += 0.5
                elif formatted_count > 0:
                    # Partial credit: proportional to how many are formatted
                    partial = 0.5 * (formatted_count / total_data)
                    print(f"PARTIAL: Component 2 — {formatted_count}/{total_data} salary values formatted ({partial:.2f} pts)")
                    total_score += partial
                else:
                    print(f"FAIL: Component 2 — No salary values are currency-formatted (0/{total_data})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Letter body salary matches first candidate's formatted salary (0.2 points)
    # The first candidate (Sarah Chen) has salary 45000, so body should show $45,000.00
    try:
        salary_para = None
        for para in doc.paragraphs:
            if 'Your annual salary will be' in para.text:
                salary_para = para.text
                break

        if salary_para is None:
            print("FAIL: Component 3 — Could not find salary paragraph")
        else:
            # Get the first data row's salary from the table
            if len(doc.tables) > 0:
                table = doc.tables[0]
                header_row = [cell.text.strip() for cell in table.rows[0].cells]
                if 'Salary' in header_row:
                    sal_col = header_row.index('Salary')
                    first_salary = table.rows[1].cells[sal_col].text.strip()

                    # Check that the body contains the same formatted value as the table
                    if CURRENCY_RE.match(first_salary) and first_salary in salary_para:
                        print(f"PASS: Component 3 — Body salary '{first_salary}' matches first candidate's table entry (0.2 pts)")
                        total_score += 0.2
                    elif '$45,000.00' in salary_para:
                        # Fallback: check against known expected value
                        print(f"PASS: Component 3 — Body contains expected '$45,000.00' (0.2 pts)")
                        total_score += 0.2
                    else:
                        print(f"FAIL: Component 3 — Body salary doesn't match first candidate. Body: {repr(salary_para)}, Table first: {first_salary}")
                else:
                    print("FAIL: Component 3 — No 'Salary' column in table")
            else:
                print("FAIL: Component 3 — No table found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
