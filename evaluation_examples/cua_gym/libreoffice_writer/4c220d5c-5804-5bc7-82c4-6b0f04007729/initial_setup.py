"""
Initial Setup: Offer letter template with unformatted Salary merge field
Task ID: writer_mt_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_018'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Company Header ---
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run("Pinnacle Technologies Inc.")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    addr_para = doc.add_paragraph()
    addr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = addr_para.add_run("2450 Innovation Drive, Suite 800\nSan Francisco, CA 94105\nTel: (415) 555-7890")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Horizontal line via bottom border on an empty paragraph
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_after = Pt(6)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(12)
    run = date_para.add_run("March 15, 2025")
    run.font.size = Pt(11)

    # --- Recipient ---
    doc.add_paragraph()  # blank line
    recip_para = doc.add_paragraph()
    run = recip_para.add_run("<CandidateName>")
    run.font.size = Pt(11)

    addr_recip = doc.add_paragraph()
    run = addr_recip.add_run("<CandidateAddress>")
    run.font.size = Pt(11)

    # --- Subject ---
    doc.add_paragraph()
    subj_para = doc.add_paragraph()
    subj_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = subj_para.add_run("Subject: Offer of Employment - <Position>")
    run.bold = True
    run.font.size = Pt(12)

    # --- Greeting ---
    doc.add_paragraph()
    greet = doc.add_paragraph()
    run = greet.add_run("Dear <CandidateName>,")
    run.font.size = Pt(11)

    # --- Body Paragraph 1 ---
    doc.add_paragraph()
    p1 = doc.add_paragraph()
    run = p1.add_run(
        "We are pleased to extend this offer of employment for the position of "
        "<Position> at Pinnacle Technologies Inc. After a thorough evaluation "
        "of your qualifications and interview performance, we believe you will "
        "be a valuable addition to our team."
    )
    run.font.size = Pt(11)

    # --- Compensation Paragraph (KEY - contains unformatted Salary merge field) ---
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    run = p2.add_run("Your annual salary will be <Salary>.")
    run.font.size = Pt(11)

    p2b = doc.add_paragraph()
    run = p2b.add_run(
        " This compensation reflects our recognition of your experience and "
        "expertise. Salary reviews are conducted annually, typically in Q1, and "
        "are based on individual performance and company results."
    )
    run.font.size = Pt(11)

    # --- Benefits Paragraph ---
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    run = p3.add_run("In addition to your base salary, you will be eligible for the following benefits:")
    run.font.size = Pt(11)

    benefits = [
        "Health, dental, and vision insurance (effective first day of employment)",
        "401(k) retirement plan with 4% company match",
        "20 days of paid time off (PTO) per year",
        "Stock options as per the company equity plan",
        "Annual performance bonus of up to 15% of base salary",
        "Professional development stipend of $2,500 per year",
    ]
    for b in benefits:
        bp = doc.add_paragraph(b, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(11)

    # --- Start Date ---
    doc.add_paragraph()
    p4 = doc.add_paragraph()
    run = p4.add_run(
        "Your anticipated start date will be <StartDate>. Please report to "
        "the Human Resources department on your first day at 9:00 AM for "
        "orientation and onboarding."
    )
    run.font.size = Pt(11)

    # --- Candidate Data Source Table ---
    doc.add_paragraph()
    table_heading = doc.add_paragraph()
    run = table_heading.add_run("Candidates Data Source")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    table_desc = doc.add_paragraph()
    run = table_desc.add_run(
        "The following table contains candidate records used for mail merge. "
        "The Salary column contains raw numeric values."
    )
    run.font.size = Pt(11)

    # Table with candidate data
    table = doc.add_table(rows=6, cols=5)
    table.style = "Table Grid"

    headers = ["CandidateName", "Position", "Salary", "StartDate", "CandidateAddress"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ["Sarah Chen", "Senior Software Engineer", "45000", "April 14, 2025", "1842 Elm Street, Oakland, CA 94602"],
        ["Marcus Johnson", "Product Manager", "72500", "May 1, 2025", "305 Pine Avenue, Apt 12, Berkeley, CA 94704"],
        ["Elena Rodriguez", "Data Scientist", "98000", "April 28, 2025", "567 Maple Court, San Jose, CA 95112"],
        ["David Kim", "UX Designer", "67500", "May 12, 2025", "1290 Cedar Lane, Palo Alto, CA 94301"],
        ["Priya Patel", "DevOps Engineer", "85000", "April 21, 2025", "428 Oak Boulevard, Fremont, CA 94536"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    # --- Closing ---
    doc.add_paragraph()
    p5 = doc.add_paragraph()
    run = p5.add_run(
        "Please confirm your acceptance of this offer by signing and returning "
        "this letter no later than March 25, 2025. Should you have any questions, "
        "please do not hesitate to contact our HR department at hr@pinnacletech.com."
    )
    run.font.size = Pt(11)

    doc.add_paragraph()
    p6 = doc.add_paragraph()
    run = p6.add_run("We look forward to welcoming you to the Pinnacle Technologies team!")
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()
    sign = doc.add_paragraph()
    run = sign.add_run("Jennifer Walsh")
    run.font.size = Pt(11)
    run.bold = True

    title_para = doc.add_paragraph()
    run = title_para.add_run("VP of Human Resources")
    run.font.size = Pt(11)

    comp_para = doc.add_paragraph()
    run = comp_para.add_run("Pinnacle Technologies Inc.")
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
