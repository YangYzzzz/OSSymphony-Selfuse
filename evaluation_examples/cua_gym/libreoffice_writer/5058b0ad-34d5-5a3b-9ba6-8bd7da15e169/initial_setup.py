"""
Initial Setup: Insert image with Through wrap and 0.3cm spacing
Task ID: writer_frd_080
Domain: libreoffice_writer

Creates a magazine article document and an illustration image.
The document has realistic content with paragraphs. No image is inserted yet.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_080'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
IMAGE_DIR = f'{WORKDIR}/Images'
IMAGE_PATH = f'{IMAGE_DIR}/illustration.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_illustration():
    """Create a simple illustration image for the task."""
    os.makedirs(IMAGE_DIR, exist_ok=True)
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new('RGB', (640, 480), color=(230, 240, 250))
    draw = ImageDraw.Draw(img)
    # Draw a simple landscape illustration
    # Sky gradient effect
    for y in range(240):
        r = int(135 + (y / 240) * 50)
        g = int(206 + (y / 240) * 30)
        b = 235
        draw.line([(0, y), (640, y)], fill=(r, g, b))
    # Ground
    draw.rectangle([0, 240, 640, 480], fill=(76, 153, 0))
    # Sun
    draw.ellipse([480, 40, 580, 140], fill=(255, 223, 0))
    # Mountains
    draw.polygon([(100, 240), (200, 120), (300, 240)], fill=(100, 100, 100))
    draw.polygon([(250, 240), (380, 80), (510, 240)], fill=(120, 120, 120))
    # Trees
    for tx in [50, 150, 400, 550]:
        draw.rectangle([tx - 5, 260, tx + 5, 320], fill=(101, 67, 33))
        draw.polygon([(tx - 30, 320), (tx, 230), (tx + 30, 320)], fill=(34, 139, 34))
    # Caption
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
    except (IOError, OSError):
        font = ImageFont.load_default()
    draw.text((200, 400), "Mountain Landscape", fill=(255, 255, 255), font=font)
    img.save(IMAGE_PATH, 'JPEG', quality=90)
    print(f'Illustration image created: {IMAGE_PATH}')


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('The Rise of Sustainable Architecture', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('How Modern Designers Are Reshaping Urban Landscapes')
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Author line
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('By Elena Vasquez | March 2025')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    author.paragraph_format.space_after = Pt(18)

    # Body paragraphs
    p1 = doc.add_paragraph()
    run = p1.add_run(
        'In cities around the world, a quiet revolution is taking place. Architects and urban '
        'planners are increasingly turning to sustainable design principles that prioritize '
        'environmental harmony alongside aesthetic beauty. From green rooftops in Copenhagen to '
        'solar-powered skyscrapers in Dubai, the movement toward eco-conscious building is '
        'transforming the way we think about the spaces we inhabit.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    run = p2.add_run(
        'Dr. Amara Okafor, a leading researcher at the Global Institute for Sustainable Design, '
        'explains that the shift is driven by both necessity and innovation. "Climate change has '
        'forced us to reconsider every aspect of construction," she notes. "But what\'s exciting '
        'is that the solutions we\'re finding are not just functional\u2014they\'re beautiful. '
        'Architects are discovering that working with nature, rather than against it, produces '
        'spaces that are more inspiring and more livable."'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # This is the paragraph where the image should be inserted
    p3 = doc.add_paragraph()
    run = p3.add_run(
        'One of the most striking examples of this trend is the Bosco Verticale in Milan, Italy. '
        'Designed by architect Stefano Boeri, the twin residential towers are home to more than '
        '900 trees and 20,000 plants. The vegetation acts as a natural air filter, absorbs CO2, '
        'and provides shade that reduces energy consumption. The project has inspired similar '
        'developments in cities from Nanjing to Eindhoven.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    p4 = doc.add_paragraph()
    run = p4.add_run(
        'Beyond individual buildings, entire neighborhoods are being reimagined with sustainability '
        'at their core. The Vauban district in Freiburg, Germany, is often cited as a model for '
        'sustainable urban living. Car-free streets, passive-energy homes, and extensive green '
        'spaces have made it one of the most desirable places to live in Europe. Residents report '
        'higher satisfaction and lower living costs compared to traditional urban developments.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    p5 = doc.add_paragraph()
    run = p5.add_run(
        'The economic case for sustainable architecture is also becoming clearer. A recent study '
        'by the World Green Building Council found that green buildings command rental premiums of '
        'up to 20 percent and have vacancy rates 4 percent lower than conventional buildings. '
        '"Sustainability is no longer a niche concern," says Marcus Lindberg, CEO of Nordic '
        'Sustainable Investments. "It\'s becoming the standard by which all new development is '
        'measured."'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    p6 = doc.add_paragraph()
    run = p6.add_run(
        'As we look to the future, the integration of technology and nature in architecture will '
        'only deepen. Innovations in bio-based materials, energy-generating facades, and smart '
        'building systems promise to make our cities not just sustainable, but regenerative\u2014'
        'actively improving the environment rather than merely reducing harm. The rise of '
        'sustainable architecture is not just a trend; it is the foundation of a new relationship '
        'between humanity and the built environment.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


# Create illustration image first, then the document
create_illustration()
create_initial()
