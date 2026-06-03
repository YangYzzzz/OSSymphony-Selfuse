"""
Initial Setup: Insert a section link to import content from changelog.docx
Task ID: writer_tech_088
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_088'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
CHANGELOG = f'{WORKDIR}/changelog.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_changelog():
    """Create the external changelog.docx file that should be linked."""
    doc = Document()

    h = doc.add_heading("Changelog", level=1)

    # Version 3.2.1
    doc.add_heading("Version 3.2.1 - 2025-11-28", level=2)
    doc.add_paragraph("Fixed critical memory leak in session handler that caused server crashes under high load.", style="List Bullet")
    doc.add_paragraph("Resolved race condition in database connection pool during failover events.", style="List Bullet")
    doc.add_paragraph("Patched XSS vulnerability in user profile display module (CVE-2025-4821).", style="List Bullet")

    # Version 3.2.0
    doc.add_heading("Version 3.2.0 - 2025-11-15", level=2)
    doc.add_paragraph("Added support for WebSocket-based real-time notifications across all dashboard views.", style="List Bullet")
    doc.add_paragraph("Implemented two-factor authentication using TOTP for administrative accounts.", style="List Bullet")
    doc.add_paragraph("Introduced dark mode theme with automatic system preference detection.", style="List Bullet")
    doc.add_paragraph("Migrated search engine from Elasticsearch 7.x to OpenSearch 2.11.", style="List Bullet")

    # Version 3.1.4
    doc.add_heading("Version 3.1.4 - 2025-10-30", level=2)
    doc.add_paragraph("Corrected pagination offset in REST API responses for collection endpoints.", style="List Bullet")
    doc.add_paragraph("Fixed file upload timeout for attachments larger than 50 MB on slow connections.", style="List Bullet")

    doc.save(CHANGELOG)
    print(f'Changelog file created: {CHANGELOG}')


def create_initial():
    """Create the main document with an empty Recent Changes section."""
    doc = Document()

    # Title
    title = doc.add_heading("Nexus Platform - Technical Documentation", level=0)

    # Overview section
    doc.add_heading("1. Overview", level=1)
    p = doc.add_paragraph(
        "The Nexus Platform is an enterprise-grade application framework designed for "
        "building scalable microservice architectures. It provides a unified API gateway, "
        "service discovery, and configuration management layer that integrates with "
        "existing CI/CD pipelines."
    )

    doc.add_paragraph(
        "This documentation covers the architecture design, deployment procedures, "
        "and operational guidelines for Nexus Platform version 3.x. Teams should "
        "refer to this document as the primary reference for all technical decisions "
        "related to platform infrastructure."
    )

    # Architecture section
    doc.add_heading("2. Architecture", level=1)
    doc.add_paragraph(
        "The platform follows a layered architecture pattern with four primary tiers:"
    )
    doc.add_paragraph("API Gateway Layer - Handles authentication, rate limiting, and request routing", style="List Bullet")
    doc.add_paragraph("Service Mesh Layer - Manages inter-service communication via Envoy sidecar proxies", style="List Bullet")
    doc.add_paragraph("Business Logic Layer - Contains domain-specific microservices written in Go and Python", style="List Bullet")
    doc.add_paragraph("Data Persistence Layer - PostgreSQL 15 for relational data, Redis 7 for caching", style="List Bullet")

    doc.add_paragraph(
        "Each service is containerized using Docker and orchestrated through Kubernetes 1.28. "
        "The cluster runs across three availability zones in AWS us-east-1 region with "
        "automatic failover configured at both the application and database levels."
    )

    # Deployment section
    doc.add_heading("3. Deployment Procedures", level=1)
    doc.add_paragraph(
        "All deployments follow the GitOps workflow managed through ArgoCD. The deployment "
        "pipeline consists of the following stages:"
    )
    doc.add_paragraph("Code review and merge to the release branch", style="List Number")
    doc.add_paragraph("Automated test suite execution (unit, integration, and E2E)", style="List Number")
    doc.add_paragraph("Container image build and vulnerability scanning via Trivy", style="List Number")
    doc.add_paragraph("Staged rollout: canary (5%) -> partial (25%) -> full deployment", style="List Number")
    doc.add_paragraph("Post-deployment health checks and metric verification", style="List Number")

    doc.add_paragraph(
        "Rollback procedures are automated. If error rates exceed 2% during canary phase, "
        "the deployment is automatically reverted within 90 seconds."
    )

    # Monitoring section
    doc.add_heading("4. Monitoring and Alerting", level=1)
    doc.add_paragraph(
        "Observability is implemented through a comprehensive stack: Prometheus for metrics "
        "collection, Grafana for dashboards, Loki for log aggregation, and Jaeger for "
        "distributed tracing. Alert rules are defined in Prometheus and routed through "
        "Alertmanager to PagerDuty for on-call engineers."
    )

    # Table of SLA metrics
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    headers = ["Service", "SLA Target", "Current Performance"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    sla_data = [
        ["API Gateway", "99.99% uptime", "99.995% (last 30 days)"],
        ["Auth Service", "99.95% uptime", "99.97% (last 30 days)"],
        ["Search Engine", "99.9% uptime", "99.92% (last 30 days)"],
        ["Notification Service", "99.9% uptime", "99.88% (last 30 days)"],
    ]
    for r, row_data in enumerate(sla_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph("")  # spacing

    # Recent Changes section - EMPTY (this is what the task asks the user to fill)
    doc.add_heading("5. Recent Changes", level=1)
    doc.add_paragraph(
        "[This section should be linked to the external changelog file for automatic updates.]"
    )

    # Contact section
    doc.add_heading("6. Contact", level=1)
    doc.add_paragraph(
        "For platform-related inquiries, reach the infrastructure team at "
        "platform-eng@nexuscorp.com or via the #nexus-platform Slack channel. "
        "On-call escalation follows the standard PagerDuty rotation schedule."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')


# Create both files
create_changelog()
create_initial()

# Launch LibreOffice Writer with the main document
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
