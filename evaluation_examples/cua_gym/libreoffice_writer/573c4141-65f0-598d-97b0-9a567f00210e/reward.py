"""
Reward Script: Change the page number format to lowercase Roman numerals (i, ii, iii).
Task ID: writer_page_016
Domain: libreoffice_writer
Scoring:
  Component 1: pgNumType fmt attribute is 'lowerRoman' (0.7 pts)
  Component 2: Footer still has centered PAGE field code (0.3 pts)
  Total: 1.0
"""

import os
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_page_016'
FILE_PATH = '/home/user/Desktop/book_preface.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — gate on this
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if not doc.sections:
        print("CRITICAL: No sections found in document.")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: Page number format changed to 'lowerRoman' (0.7 points)
    # The task requires changing pgNumType fmt from 'decimal' to 'lowerRoman'.
    # This is stored in w:sectPr/w:pgNumType/@w:fmt in the section properties XML.
    try:
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(f'{{{W_NS}}}pgNumType')
        if pg_num_type is not None:
            fmt_value = pg_num_type.get(f'{{{W_NS}}}fmt')
            if fmt_value == 'lowerRoman':
                print(f"PASS: Component 1 — pgNumType fmt='lowerRoman' (0.7 pts)")
                total_score += 0.7
            else:
                print(f"FAIL: Component 1 — expected pgNumType fmt='lowerRoman', found: {repr(fmt_value)}")
        else:
            print(f"FAIL: Component 1 — pgNumType element not found in section properties")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Footer still has centered PAGE field code (0.3 points)
    # Verifies that the footer is not broken: it still contains a PAGE field code
    # (the field instruction ' PAGE ') AND the paragraph alignment is CENTER.
    # This FAILS on initial_env because Component 1 gates the whole scoring
    # (initial has 'decimal', not 'lowerRoman'), but we want to independently
    # check alignment + field code presence here as a completeness check.
    # Note: We only award this component when Component 1 also passes, ensuring
    # it cannot score on the initial artifact.
    try:
        footer = section.footer
        page_field_count = 0
        centered_para_count = 0

        if footer and footer.paragraphs:
            for p in footer.paragraphs:
                # Check for PAGE instrText field
                para_xml = p._element.xml
                tree = ET.fromstring(para_xml)
                for instr in tree.findall(f'.//{{{W_NS}}}instrText'):
                    if instr.text and 'PAGE' in instr.text:
                        page_field_count += 1
                # Check alignment
                if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    centered_para_count += 1

        footer_has_page_field = page_field_count > 0
        footer_is_centered = centered_para_count > 0

        # Only award this component if Component 1 also passed (score >= 0.7)
        # to avoid awarding partial credit on initial env state
        if total_score >= 0.7:
            if footer_has_page_field and footer_is_centered:
                print(f"PASS: Component 2 — footer has centered PAGE field code (0.3 pts)")
                total_score += 0.3
            elif not footer_has_page_field:
                print(f"FAIL: Component 2 — footer PAGE field code not found")
            elif not footer_is_centered:
                print(f"FAIL: Component 2 — footer paragraph not centered (alignment={footer.paragraphs[0].paragraph_format.alignment if footer.paragraphs else 'none'})")
        else:
            print(f"SKIP: Component 2 — skipped because Component 1 failed (format not changed)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
