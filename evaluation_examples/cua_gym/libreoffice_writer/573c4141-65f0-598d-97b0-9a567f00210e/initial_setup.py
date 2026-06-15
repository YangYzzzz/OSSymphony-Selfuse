"""
Initial Setup: Book preface document with Arabic page numbers in footer
Task ID: writer_page_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

WORKDIR = '/home/user'  # VM path
DESKTOP = '/home/user/Desktop'
TASK_ID = 'book_preface'
OUTPUT = f'{DESKTOP}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph, numFmt='decimal'):
    """Add a page number field to the paragraph with specified format.
    numFmt: 'decimal' for Arabic (1,2,3), 'lowerRoman' for roman (i,ii,iii)
    """
    # Clear existing content
    for run in paragraph.runs:
        run.text = ''

    # We add the PAGE field directly
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar_end)


def set_page_number_format(section, fmt='decimal', start=1):
    """Set the page number format in section properties.
    fmt: 'decimal' for Arabic, 'lowerRoman' for Roman, etc.
    """
    sectPr = section._sectPr
    # Remove existing pgNumType if present
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    # Add new pgNumType
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Configure section: A4, portrait, specific margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4 width
    section.page_height = Cm(29.7)  # A4 height
    # Margins: top=2.54cm, bottom=2.54cm, left=3.0cm, right=2.0cm
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # Set page number format to Arabic (decimal) — task is to change to Roman
    set_page_number_format(section, fmt='decimal', start=1)

    # --- Footer with centered page number ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Clear any default content
    fp.clear()
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp, numFmt='decimal')

    # ---- Document content: 4-page book preface ----

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = title.add_run("PREFACE")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.name = "Times New Roman"

    doc.add_paragraph()  # spacing

    # Opening paragraph
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p1.paragraph_format.first_line_indent = Cm(1.25)
    run1 = p1.add_run(
        "This book was born out of years of careful observation and a deep curiosity "
        "about the natural world. When I first began collecting notes for what would "
        "eventually become this volume, I had no clear sense of where the work would "
        "lead. The journey from those scattered observations to a coherent manuscript "
        "has been long and often surprising."
    )
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2.paragraph_format.first_line_indent = Cm(1.25)
    run2 = p2.add_run(
        "The core ideas explored here draw on a wide range of sources: field notes "
        "gathered over two decades, conversations with colleagues at universities and "
        "research stations around the world, and a careful re-reading of foundational "
        "texts that shaped the discipline. I have tried, wherever possible, to let "
        "the evidence speak for itself rather than imposing an overly rigid theoretical "
        "framework on observations that remain, in many respects, wonderfully complex."
    )
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p3.paragraph_format.first_line_indent = Cm(1.25)
    run3 = p3.add_run(
        "Chapter One introduces the key conceptual distinctions that underpin the "
        "entire analysis. Readers who are already familiar with the existing literature "
        "may find that some of these definitions overlap with work they already know; "
        "I have tried to be explicit about points of contact and divergence. For those "
        "approaching the subject for the first time, I hope the opening chapter "
        "provides a sufficiently clear foundation without sacrificing intellectual rigor."
    )
    run3.font.name = "Times New Roman"
    run3.font.size = Pt(12)

    # Page break for page 2
    doc.add_page_break()

    # Section heading for page 2
    h2 = doc.add_paragraph()
    h2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_h2 = h2.add_run("Sources and Methods")
    run_h2.bold = True
    run_h2.font.size = Pt(14)
    run_h2.font.name = "Times New Roman"

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p4.paragraph_format.first_line_indent = Cm(1.25)
    run4 = p4.add_run(
        "The primary sources consulted for this work are held in several archives, "
        "including the National Library collection and the private papers of Professor "
        "Eleanor Harwick, whose unpublished correspondence proved invaluable. Secondary "
        "sources are listed in full at the end of the volume. Where I have drawn on "
        "unpublished materials, I have sought permission from the relevant institutions "
        "and individuals, and I am grateful to all who responded generously."
    )
    run4.font.name = "Times New Roman"
    run4.font.size = Pt(12)

    p5 = doc.add_paragraph()
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p5.paragraph_format.first_line_indent = Cm(1.25)
    run5 = p5.add_run(
        "The quantitative data presented in Chapters Three through Six were collected "
        "using standardized protocols developed in collaboration with research teams at "
        "the University of Thessaloniki and the Coastal Research Station in Bergen. "
        "Statistical analyses were conducted using industry-standard software, and all "
        "code used for data processing has been made available in a public repository. "
        "Readers wishing to replicate or extend the analyses will find full documentation "
        "in Appendix B."
    )
    run5.font.name = "Times New Roman"
    run5.font.size = Pt(12)

    p6 = doc.add_paragraph()
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p6.paragraph_format.first_line_indent = Cm(1.25)
    run6 = p6.add_run(
        "Field observations were recorded using a combination of written notes, "
        "photographic documentation, and digital audio recordings. All recordings were "
        "subsequently transcribed and coded by at least two independent researchers. "
        "Disagreements were resolved through discussion and, where necessary, by "
        "reference to the original recordings. This multi-stage process was time-consuming "
        "but essential for ensuring the reliability of the data."
    )
    run6.font.name = "Times New Roman"
    run6.font.size = Pt(12)

    # Page break for page 3
    doc.add_page_break()

    # Section heading for page 3
    h3 = doc.add_paragraph()
    h3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_h3 = h3.add_run("Acknowledgements")
    run_h3.bold = True
    run_h3.font.size = Pt(14)
    run_h3.font.name = "Times New Roman"

    doc.add_paragraph()

    p7 = doc.add_paragraph()
    p7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p7.paragraph_format.first_line_indent = Cm(1.25)
    run7 = p7.add_run(
        "Many people contributed to the development of this book, and it is a pleasure "
        "to acknowledge their help here. My colleagues at the Institute for Applied "
        "Research have been a constant source of stimulation and support. I am "
        "particularly grateful to Dr. Miriam Osei-Bonsu and Professor David Nakamura, "
        "both of whom read earlier drafts with extraordinary care and offered criticisms "
        "that materially improved the final text."
    )
    run7.font.name = "Times New Roman"
    run7.font.size = Pt(12)

    p8 = doc.add_paragraph()
    p8.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p8.paragraph_format.first_line_indent = Cm(1.25)
    run8 = p8.add_run(
        "The editorial team at the Press have been patient and professional throughout "
        "a process that took rather longer than any of us anticipated. In particular, "
        "I wish to thank Senior Editor Clare Whitfield, whose gentle but persistent "
        "encouragement kept the project moving during several difficult periods. "
        "The anonymous reviewers commissioned by the Press provided rigorous and "
        "constructive assessments that significantly strengthened the manuscript."
    )
    run8.font.name = "Times New Roman"
    run8.font.size = Pt(12)

    p9 = doc.add_paragraph()
    p9.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p9.paragraph_format.first_line_indent = Cm(1.25)
    run9 = p9.add_run(
        "My research assistants — Tomás Rivera, Yuki Watanabe, and Priya Krishnamurthy "
        "— carried out a great deal of the painstaking work of checking references, "
        "assembling the index, and preparing the figures. Without their diligence, the "
        "final manuscript would contain many more errors than it does. Those that remain "
        "are, of course, entirely my own responsibility."
    )
    run9.font.name = "Times New Roman"
    run9.font.size = Pt(12)

    # Page break for page 4
    doc.add_page_break()

    # Section heading for page 4
    h4 = doc.add_paragraph()
    h4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_h4 = h4.add_run("A Note on Terminology")
    run_h4.bold = True
    run_h4.font.size = Pt(14)
    run_h4.font.name = "Times New Roman"

    doc.add_paragraph()

    p10 = doc.add_paragraph()
    p10.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p10.paragraph_format.first_line_indent = Cm(1.25)
    run10 = p10.add_run(
        "The terminology used in this field is unfortunately not fully standardized, "
        "and different researchers sometimes use the same term to refer to subtly "
        "different phenomena, or different terms to refer to the same thing. Where "
        "possible, I have followed the conventions established in the landmark study "
        "by Harwick and Patel (2009), which remains the most widely cited reference "
        "work in the area."
    )
    run10.font.name = "Times New Roman"
    run10.font.size = Pt(12)

    p11 = doc.add_paragraph()
    p11.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p11.paragraph_format.first_line_indent = Cm(1.25)
    run11 = p11.add_run(
        "In cases where I have departed from that convention, or where Harwick and "
        "Patel's usage itself is contested, I have explained my choice in footnotes. "
        "Readers who are primarily interested in the empirical findings, rather than "
        "the terminological debates, may safely skip those notes without loss of "
        "continuity. The main text is designed to be readable without specialist "
        "knowledge of the definitional controversies."
    )
    run11.font.name = "Times New Roman"
    run11.font.size = Pt(12)

    p12 = doc.add_paragraph()
    p12.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p12.paragraph_format.first_line_indent = Cm(1.25)
    run12 = p12.add_run(
        "Finally, a word about scope. This book is not intended as a comprehensive "
        "survey of the field. Such surveys exist, and some are excellent; readers "
        "seeking an overview are directed to the works listed in the bibliographic "
        "note at the end of Chapter One. My aim here is more focused: to present "
        "a sustained argument about a specific and contested question, and to do so "
        "with sufficient rigor that the argument can be tested and, if necessary, "
        "refuted by future research.\n\nThe Author\nMarch 2025"
    )
    run12.font.name = "Times New Roman"
    run12.font.size = Pt(12)

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
