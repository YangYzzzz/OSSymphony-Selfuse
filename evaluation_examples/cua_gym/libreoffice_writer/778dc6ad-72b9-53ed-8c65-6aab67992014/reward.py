"""
Reward Script: Set cell padding (margins) for all cells in table
Task ID: writer_tbl_032
Domain: libreoffice_writer
Scoring:
  Component 1: All 12 cells have explicit tcMar (cell margin) element set    (0.30 pts)
  Component 2: Top and bottom margins are 113 dxa (0.2 cm) for all cells     (0.35 pts)
  Component 3: Left and right margins are 170 dxa (0.3 cm) for all cells     (0.35 pts)
  Total: 1.0

Conversion: 1 cm = 567 twips (dxa)
  0.2 cm = 113.4 -> rounded to 113 dxa
  0.3 cm = 170.1 -> rounded to 170 dxa

Initial env: No cell margins set anywhere (defaults).
Golden env:  All 12 cells have per-cell tcMar with exact dxa values above.
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_032'

# Expected margin values in dxa (twips)
# 0.2 cm * 567 dxa/cm = 113.4 -> 113 dxa
# 0.3 cm * 567 dxa/cm = 170.1 -> 170 dxa
EXPECTED_TOP_BOTTOM_DXA = 113   # 0.2 cm
EXPECTED_LEFT_RIGHT_DXA = 170   # 0.3 cm
TOLERANCE_DXA = 5               # allow +/-5 twips tolerance (~0.009 cm)

EXPECTED_CONTENTS = [
    ['Item', 'Qty', 'Unit Price'],
    ['Pencils', '100', '$0.50'],
    ['Notebooks', '50', '$2.00'],
    ['Erasers', '200', '$0.25'],
]


def verify_task(file_path):
    """
    Verify that all cells in the table have cell margins set to:
      top=0.2cm, bottom=0.2cm, left=0.3cm, right=0.3cm
    Returns a float between 0.0 and 1.0.
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print('CRITICAL: Cannot load file %s: %s' % (file_path, e))
        print('REWARD: 0.0')
        return 0.0

    # Precondition: exactly 1 table present
    if len(doc.tables) == 0:
        print('CRITICAL: No table found in document.')
        print('REWARD: 0.0')
        return 0.0

    table = doc.tables[0]

    # Precondition: table has 4 rows and 3 columns
    if len(table.rows) != 4 or len(table.columns) != 3:
        print('CRITICAL: Table dimensions wrong: %d rows, %d cols (expected 4x3).' % (
            len(table.rows), len(table.columns)))
        print('REWARD: 0.0')
        return 0.0

    # Precondition gate: cell contents unchanged
    contents_ok = True
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            actual = cell.text.strip()
            expected = EXPECTED_CONTENTS[i][j]
            if actual != expected:
                print('PRECONDITION FAIL: Cell [%d,%d] content changed. Expected "%s", found "%s".' % (
                    i, j, expected, actual))
                contents_ok = False
    if not contents_ok:
        print('CRITICAL: Cell content precondition failed.')
        print('REWARD: 0.0')
        return 0.0

    print('PRECONDITION PASS: Table structure and cell contents intact (4 rows x 3 cols).')

    # Collect per-cell margin data
    cell_margins = {}  # (i,j) -> {'top': int, 'bottom': int, 'left': int, 'right': int}
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is None:
                cell_margins[(i, j)] = None
                continue
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is None:
                cell_margins[(i, j)] = None
                continue
            margins = {}
            for side in ('top', 'bottom', 'left', 'right'):
                elem = tcMar.find(qn('w:' + side))
                if elem is not None:
                    w_val = elem.get(qn('w:w'))
                    try:
                        margins[side] = int(w_val)
                    except (TypeError, ValueError):
                        margins[side] = None
                else:
                    margins[side] = None
            cell_margins[(i, j)] = margins

    # Component 1: All 12 cells have explicit tcMar element (0.30 pts)
    # This is the key indicator that margins were explicitly set at all.
    # Fails on initial (no tcMar anywhere), passes on golden.
    try:
        cells_with_margins = sum(1 for v in cell_margins.values() if v is not None)
        total_cells = 4 * 3  # 12 cells
        if cells_with_margins == total_cells:
            print('PASS: Component 1 — All %d cells have explicit tcMar set. (0.30 pts)' % total_cells)
            total_score += 0.30
        else:
            print('FAIL: Component 1 — Only %d/%d cells have tcMar set.' % (cells_with_margins, total_cells))
    except Exception as e:
        print('ERROR: Component 1 — %s' % e)

    # Component 2: Top and bottom margins are 113 dxa (0.2 cm) for all cells (0.35 pts)
    try:
        top_bottom_correct = 0
        top_bottom_failed = []
        for (i, j), margins in cell_margins.items():
            if margins is None:
                top_bottom_failed.append('[%d,%d]: no tcMar' % (i, j))
                continue
            top_ok = margins.get('top') is not None and abs(margins['top'] - EXPECTED_TOP_BOTTOM_DXA) <= TOLERANCE_DXA
            bot_ok = margins.get('bottom') is not None and abs(margins['bottom'] - EXPECTED_TOP_BOTTOM_DXA) <= TOLERANCE_DXA
            if top_ok and bot_ok:
                top_bottom_correct += 1
            else:
                top_bottom_failed.append('[%d,%d]: top=%s, bottom=%s (expected ~%d)' % (
                    i, j, margins.get('top'), margins.get('bottom'), EXPECTED_TOP_BOTTOM_DXA))
        if top_bottom_correct == total_cells:
            print('PASS: Component 2 — All %d cells have top/bottom margin = %d dxa (0.2 cm). (0.35 pts)' % (
                total_cells, EXPECTED_TOP_BOTTOM_DXA))
            total_score += 0.35
        else:
            print('FAIL: Component 2 — %d/%d cells have correct top/bottom margins.' % (
                top_bottom_correct, total_cells))
            for detail in top_bottom_failed[:4]:
                print('  ' + detail)
    except Exception as e:
        print('ERROR: Component 2 — %s' % e)

    # Component 3: Left and right margins are 170 dxa (0.3 cm) for all cells (0.35 pts)
    try:
        left_right_correct = 0
        left_right_failed = []
        for (i, j), margins in cell_margins.items():
            if margins is None:
                left_right_failed.append('[%d,%d]: no tcMar' % (i, j))
                continue
            left_ok = margins.get('left') is not None and abs(margins['left'] - EXPECTED_LEFT_RIGHT_DXA) <= TOLERANCE_DXA
            right_ok = margins.get('right') is not None and abs(margins['right'] - EXPECTED_LEFT_RIGHT_DXA) <= TOLERANCE_DXA
            if left_ok and right_ok:
                left_right_correct += 1
            else:
                left_right_failed.append('[%d,%d]: left=%s, right=%s (expected ~%d)' % (
                    i, j, margins.get('left'), margins.get('right'), EXPECTED_LEFT_RIGHT_DXA))
        if left_right_correct == total_cells:
            print('PASS: Component 3 — All %d cells have left/right margin = %d dxa (0.3 cm). (0.35 pts)' % (
                total_cells, EXPECTED_LEFT_RIGHT_DXA))
            total_score += 0.35
        else:
            print('FAIL: Component 3 — %d/%d cells have correct left/right margins.' % (
                left_right_correct, total_cells))
            for detail in left_right_failed[:4]:
                print('  ' + detail)
    except Exception as e:
        print('ERROR: Component 3 — %s' % e)

    final_score = min(total_score, 1.0)
    print()
    print('Score: %.2f/1.0' % total_score)
    print('REWARD: %.1f' % final_score)
    return final_score


# Default: test against canonical artifact path on VM Desktop
file_path = '/home/user/Desktop/padded_table.docx'
if not os.path.exists(file_path):
    print('File not found: %s' % file_path)
    print('REWARD: 0.0')
else:
    verify_task(file_path)
