"""
Initial Setup: Table cell padding task - create a docx with a 4x3 table with default cell margins
Task ID: writer_tbl_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, Cm

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_032'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/padded_table.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Add a title paragraph for context
    doc.add_heading('Office Supplies Order', level=1)

    # Create a 4-row x 3-column table with default cell padding
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Row 1: Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Item'
    header_cells[1].text = 'Qty'
    header_cells[2].text = 'Unit Price'

    # Make header row bold
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Row 2
    row2 = table.rows[1].cells
    row2[0].text = 'Pencils'
    row2[1].text = '100'
    row2[2].text = '$0.50'

    # Row 3
    row3 = table.rows[2].cells
    row3[0].text = 'Notebooks'
    row3[1].text = '50'
    row3[2].text = '$2.00'

    # Row 4
    row4 = table.rows[3].cells
    row4[0].text = 'Erasers'
    row4[1].text = '200'
    row4[2].text = '$0.25'

    # NOTE: Do NOT set any custom cell margins here — task asks agent to set them
    # Default cell margins are used (no tblCellMar or tcMar XML elements set)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
