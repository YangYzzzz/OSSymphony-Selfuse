"""
Reward Script: Community Garden Spring Planting Day Flyer
Task ID: writer_creative_030
Domain: libreoffice_writer
Scoring:
  - Component 1: Title formatting (32pt, bold, center, green #228B22) — 0.35 pts
  - Component 2: Date/Time/Location formatting — 0.20 pts
  - Component 3: Plot table exists with correct structure and content — 0.25 pts
  - Component 4: Table header bold + background #90EE90 — 0.10 pts
  - Component 5: 'Bring...' line (14pt, bold, center) + 'Contact...' line (12pt, center) — 0.10 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_030'
FILE_PATH = '/home/user/Desktop/garden_planting_flyer.docx'


def color_distance(rgb_obj, target_rgb):
    """Calculate Euclidean distance between two RGB colors."""
    if rgb_obj is None:
        return 999
    r1, g1, b1 = rgb_obj[0], rgb_obj[1], rgb_obj[2]
    r2, g2, b2 = target_rgb
    return sqrt((r1 - r2)**2 + (g1 - g2)**2 + (b1 - b2)**2)


def find_paragraph_by_text(doc, text_substr):
    """Find first paragraph whose text contains the given substring."""
    for para in doc.paragraphs:
        if text_substr in para.text:
            return para
    return None


def get_run_size_pt(run):
    """Get run font size in points, or None."""
    if run.font.size:
        return run.font.size.pt
    return None


def get_run_color(run):
    """Get RGBColor from run, or None."""
    try:
        if run.font.color and run.font.color.type is not None:
            return run.font.color.rgb
    except Exception:
        pass
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Title 'Spring Planting Day!' formatting (0.35 points)
    # Expected: font_size=32pt, bold=True, alignment=CENTER, color=#228B22
    # This FAILS on initial (12pt, not bold, LEFT, no color) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        title_para = find_paragraph_by_text(doc, 'Spring Planting Day!')
        if title_para is None:
            print("FAIL: Component 1 — Title paragraph 'Spring Planting Day!' not found")
        else:
            title_score = 0.0
            runs = [r for r in title_para.runs if r.text.strip()]

            # Check size: 32pt
            size_ok = any(get_run_size_pt(r) == 32.0 for r in runs)
            # Check bold
            bold_ok = any(r.font.bold is True for r in runs)
            # Check alignment CENTER
            align_ok = (title_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            # Check color: #228B22 (forest green), use tolerance of 20
            GREEN_228B22 = (0x22, 0x8B, 0x22)
            color_ok = any(
                color_distance(get_run_color(r), GREEN_228B22) < 20
                for r in runs
            )

            checks = [size_ok, bold_ok, align_ok, color_ok]
            passed = sum(checks)
            partial = passed / 4.0 * 0.35

            details = (
                f"size_32pt={size_ok}, bold={bold_ok}, "
                f"align_center={align_ok}, color_green={color_ok}"
            )
            if passed == 4:
                print(f"PASS: Component 1 — Title fully formatted ({details}) (0.35 pts)")
                total_score += 0.35
            elif passed > 0:
                print(f"PARTIAL: Component 1 — Title partial ({details}) ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 1 — Title not formatted ({details})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Date/Time/Location formatting (0.20 points)
    # Date: 22pt, bold, CENTER; Time: 18pt, CENTER; Location: 14pt, CENTER
    # This FAILS on initial (all 12pt, LEFT, not bold) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        sub_score = 0.0

        # Date paragraph: 'Saturday, April 12, 2026' → 22pt, bold, CENTER
        date_para = find_paragraph_by_text(doc, 'Saturday, April 12, 2026')
        if date_para is not None:
            d_runs = [r for r in date_para.runs if r.text.strip()]
            date_size_ok = any(get_run_size_pt(r) == 22.0 for r in d_runs)
            date_bold_ok = any(r.font.bold is True for r in d_runs)
            date_align_ok = (date_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            if date_size_ok and date_bold_ok and date_align_ok:
                print("PASS: Component 2a — Date: 22pt, bold, center")
                sub_score += 0.07
            else:
                print(f"FAIL: Component 2a — Date: size22={date_size_ok}, bold={date_bold_ok}, center={date_align_ok}")

        # Time paragraph: '9:00 AM - 1:00 PM' → 18pt, CENTER
        time_para = find_paragraph_by_text(doc, '9:00 AM')
        if time_para is not None:
            t_runs = [r for r in time_para.runs if r.text.strip()]
            time_size_ok = any(get_run_size_pt(r) == 18.0 for r in t_runs)
            time_align_ok = (time_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            if time_size_ok and time_align_ok:
                print("PASS: Component 2b — Time: 18pt, center")
                sub_score += 0.06
            else:
                print(f"FAIL: Component 2b — Time: size18={time_size_ok}, center={time_align_ok}")

        # Location paragraphs: 14pt, CENTER
        loc1_para = find_paragraph_by_text(doc, 'Riverside Community Garden')
        loc2_para = find_paragraph_by_text(doc, '450 River Road')
        loc_ok_count = 0
        for lp in [loc1_para, loc2_para]:
            if lp is not None:
                l_runs = [r for r in lp.runs if r.text.strip()]
                l_size_ok = any(get_run_size_pt(r) == 14.0 for r in l_runs)
                l_align_ok = (lp.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                if l_size_ok and l_align_ok:
                    loc_ok_count += 1
        if loc_ok_count == 2:
            print("PASS: Component 2c — Both location lines: 14pt, center")
            sub_score += 0.07
        elif loc_ok_count == 1:
            print("PARTIAL: Component 2c — One location line: 14pt, center")
            sub_score += 0.03

        total_score += sub_score
        print(f"Component 2 total: {sub_score:.2f}/0.20")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Plot table with correct structure and content (0.25 points)
    # Expected: 1 table, 3 columns, 6 rows (1 header + 5 data)
    # Header: 'Plot', 'Size', 'Sun Exposure'
    # Data rows: A1/10x10ft/Full Sun, A2/10x10ft/Full Sun,
    #            B1/8x12ft/Partial Shade, B2/8x12ft/Partial Shade, C1/6x8ft/Full Sun
    # This FAILS on initial (no table) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No tables found in document")
        else:
            table = doc.tables[0]
            rows = table.rows
            cols = table.columns

            # Check structure
            has_correct_cols = len(cols) == 3
            has_correct_rows = len(rows) == 6  # 1 header + 5 data rows

            # Check header content
            expected_headers = ['Plot', 'Size', 'Sun Exposure']
            header_cells = [c.text.strip() for c in rows[0].cells]
            headers_ok = all(
                any(eh.lower() in hc.lower() for hc in header_cells)
                for eh in expected_headers
            )

            # Check data rows content
            expected_data = [
                ('A1', '10x10', 'Full Sun'),
                ('A2', '10x10', 'Full Sun'),
                ('B1', '8x12', 'Partial Shade'),
                ('B2', '8x12', 'Partial Shade'),
                ('C1', '6x8', 'Full Sun'),
            ]
            data_rows_fail_count = 0
            for ri, (plot, size, sun) in enumerate(expected_data):
                if ri + 1 < len(rows):
                    row_texts = [c.text.strip() for c in rows[ri + 1].cells]
                    row_text_joined = ' '.join(row_texts)
                    if plot not in row_text_joined or size not in row_text_joined or sun not in row_text_joined:
                        data_rows_fail_count += 1
                        print(f"FAIL: Component 3 data row {ri+1}: expected {(plot, size, sun)}, found {row_texts}")
                else:
                    data_rows_fail_count += 1
            data_rows_ok = (data_rows_fail_count == 0)

            sub_score_3 = 0.0
            if has_correct_cols and has_correct_rows:
                print(f"PASS: Component 3a — Table structure: {len(rows)} rows x {len(cols)} cols")
                sub_score_3 += 0.10
            else:
                print(f"FAIL: Component 3a — Table structure: {len(rows)} rows x {len(cols)} cols (expected 6x3)")

            if headers_ok:
                print(f"PASS: Component 3b — Table headers correct: {header_cells}")
                sub_score_3 += 0.07
            else:
                print(f"FAIL: Component 3b — Table headers: found {header_cells}, expected {expected_headers}")

            if data_rows_ok:
                print("PASS: Component 3c — All 5 data rows present and correct")
                sub_score_3 += 0.08
            else:
                print("FAIL: Component 3c — Data rows incomplete or incorrect")

            total_score += sub_score_3
            print(f"Component 3 total: {sub_score_3:.2f}/0.25")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Table header formatting — bold + background #90EE90 (0.10 points)
    # This FAILS on initial (no table) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        if len(doc.tables) > 0:
            table = doc.tables[0]
            header_row = table.rows[0]

            LIGHT_GREEN = (0x90, 0xEE, 0x90)
            header_bold_count = 0
            header_bg_count = 0

            for cell in header_row.cells:
                # Check bold
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.bold is True:
                            header_bold_count += 1
                            break

                # Check background color via XML
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill_hex = shd.get(qn('w:fill'), '')
                        if fill_hex:
                            try:
                                fill_rgb = (
                                    int(fill_hex[0:2], 16),
                                    int(fill_hex[2:4], 16),
                                    int(fill_hex[4:6], 16),
                                )
                                if color_distance(fill_rgb, LIGHT_GREEN) < 30:
                                    header_bg_count += 1
                            except Exception:
                                pass

            sub_score_4 = 0.0
            n_header_cells = len(header_row.cells)
            if header_bold_count == n_header_cells:
                print(f"PASS: Component 4a — Header cells bold ({header_bold_count}/{n_header_cells})")
                sub_score_4 += 0.05
            else:
                print(f"FAIL: Component 4a — Header bold: {header_bold_count}/{n_header_cells} cells bold")

            if header_bg_count == n_header_cells:
                print(f"PASS: Component 4b — Header cells background #90EE90 ({header_bg_count}/{n_header_cells})")
                sub_score_4 += 0.05
            else:
                print(f"FAIL: Component 4b — Header bg: {header_bg_count}/{n_header_cells} cells have light green bg")

            total_score += sub_score_4
            print(f"Component 4 total: {sub_score_4:.2f}/0.10")
        else:
            print("FAIL: Component 4 — No table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: 'Bring...' line (14pt, bold, center) + 'Contact...' (12pt, center) (0.10 pts)
    # This FAILS on initial (12pt, not bold, LEFT) → PASSES on golden
    # -------------------------------------------------------------------------
    try:
        sub_score_5 = 0.0

        # 'Bring your own seeds...' → 14pt, bold, center
        bring_para = find_paragraph_by_text(doc, 'Bring your own seeds')
        if bring_para is not None:
            b_runs = [r for r in bring_para.runs if r.text.strip()]
            b_size_ok = any(get_run_size_pt(r) == 14.0 for r in b_runs)
            b_bold_ok = any(r.font.bold is True for r in b_runs)
            b_align_ok = (bring_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            if b_size_ok and b_bold_ok and b_align_ok:
                print(f"PASS: Component 5a — 'Bring...' line: 14pt, bold, center")
                sub_score_5 += 0.05
            else:
                print(f"FAIL: Component 5a — 'Bring...' size14={b_size_ok}, bold={b_bold_ok}, center={b_align_ok}")
        else:
            print("FAIL: Component 5a — 'Bring your own seeds...' paragraph not found")

        # 'Contact: Maria...' → 12pt, center
        contact_para = find_paragraph_by_text(doc, 'Contact:')
        if contact_para is not None:
            c_runs = [r for r in contact_para.runs if r.text.strip()]
            # 12pt can also be default (None), accept None or 12.0
            c_size_ok = all(
                (get_run_size_pt(r) is None or get_run_size_pt(r) == 12.0)
                for r in c_runs
            ) if c_runs else False
            c_align_ok = (contact_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            if c_align_ok:
                print(f"PASS: Component 5b — 'Contact...' line: center aligned")
                sub_score_5 += 0.05
            else:
                print(f"FAIL: Component 5b — 'Contact...' align_center={c_align_ok}")
        else:
            print("FAIL: Component 5b — 'Contact:' paragraph not found")

        total_score += sub_score_5
        print(f"Component 5 total: {sub_score_5:.2f}/0.10")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -------------------------------------------------------------------------
    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
