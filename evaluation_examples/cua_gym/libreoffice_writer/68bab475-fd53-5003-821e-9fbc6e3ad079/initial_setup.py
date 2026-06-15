"""
Initial Setup: Community garden spring planting day flyer
Task ID: writer_creative_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_creative_030'
OUTPUT = f'{WORKDIR}/Desktop/garden_planting_flyer.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # All text: 12pt Times New Roman, left-aligned (no bold, no colors, no centering)
    def add_plain_para(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = False
        run.italic = False
        # Default alignment is left
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return para

    # Title (plain — agent will make it big, bold, green, centered)
    add_plain_para('Spring Planting Day!')

    # Date and time
    add_plain_para('Saturday, April 12, 2026')
    add_plain_para('9:00 AM - 1:00 PM')

    # Location
    add_plain_para('Riverside Community Garden')
    add_plain_para('450 River Road, Eugene, OR')

    # Available Plots heading (plain — agent will bold and resize)
    add_plain_para('Available Plots:')

    # Individual plot listings as plain paragraphs (agent will convert to table)
    add_plain_para('Plot A1 - 10x10 ft - Full Sun')
    add_plain_para('Plot A2 - 10x10 ft - Full Sun')
    add_plain_para('Plot B1 - 8x12 ft - Partial Shade')
    add_plain_para('Plot B2 - 8x12 ft - Partial Shade')
    add_plain_para('Plot C1 - 6x8 ft - Full Sun')

    # Call to action (plain — agent will bold, resize, center)
    add_plain_para('Bring your own seeds, gloves, and tools!')

    # Contact
    add_plain_para('Contact: Maria at (541) 555-0189')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
