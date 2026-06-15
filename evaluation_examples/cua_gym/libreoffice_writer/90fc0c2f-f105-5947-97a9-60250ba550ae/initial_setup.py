"""
Initial Setup: expense_report.docx with expense table (C6 empty, no formula)
Task ID: writer_tbl_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_028'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/expense_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Add a title paragraph
    title = doc.add_heading('Expense Report', level=1)

    # Create the expense table: 6 rows x 3 columns
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    # Row 1 (header): Category | Description | Amount
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Category'
    header_cells[1].text = 'Description'
    header_cells[2].text = 'Amount'

    # Make header bold
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Row 2: Travel | Flight tickets | 450
    row2 = table.rows[1].cells
    row2[0].text = 'Travel'
    row2[1].text = 'Flight tickets'
    row2[2].text = '450'

    # Row 3: Hotel | 3 nights | 360
    row3 = table.rows[2].cells
    row3[0].text = 'Hotel'
    row3[1].text = '3 nights'
    row3[2].text = '360'

    # Row 4: Meals | Per diem | 120
    row4 = table.rows[3].cells
    row4[0].text = 'Meals'
    row4[1].text = 'Per diem'
    row4[2].text = '120'

    # Row 5: Transport | Taxi and rental | 95
    row5 = table.rows[4].cells
    row5[0].text = 'Transport'
    row5[1].text = 'Taxi and rental'
    row5[2].text = '95'

    # Row 6: Total | (empty) | (empty - no formula yet, that is the task)
    row6 = table.rows[5].cells
    row6[0].text = 'Total'
    row6[1].text = ''
    row6[2].text = ''  # C6 must be empty; the agent will add the SUM formula

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
