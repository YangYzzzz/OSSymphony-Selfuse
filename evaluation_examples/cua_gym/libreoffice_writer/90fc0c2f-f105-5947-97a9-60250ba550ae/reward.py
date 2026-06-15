"""
Reward Script: Add a SUM formula in the last cell of column 3 in the expense_report table.
Task ID: writer_tbl_028
Domain: libreoffice_writer
Scoring:
  Component 1: Cell [5,2] (row 6, col 3) displays the computed total '1025'  — 0.5 pts
  Component 2: Cell [5,2] contains a Writer field formula with SUM            — 0.5 pts
  Data integrity (other cells unchanged): used as a precondition gate only,
    not as a scoring component (would pass on initial_env too).
Total: 1.0
"""

import os
from lxml import etree
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_FILE = 'expense_report.docx'

# Namespace for OOXML
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Expected data for precondition gate (rows 0-4, all columns + row 5 col 0 and col 1)
# These cells must be intact for the task to be meaningful.
PRECONDITION_CELLS = {
    (0, 0): 'Category',
    (0, 1): 'Description',
    (0, 2): 'Amount',
    (1, 0): 'Travel',
    (1, 1): 'Flight tickets',
    (1, 2): '450',
    (2, 0): 'Hotel',
    (2, 1): '3 nights',
    (2, 2): '360',
    (3, 0): 'Meals',
    (3, 1): 'Per diem',
    (3, 2): '120',
    (4, 0): 'Transport',
    (4, 1): 'Taxi and rental',
    (4, 2): '95',
    (5, 0): 'Total',
    (5, 1): '',
}

EXPECTED_TOTAL = '1025'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — gate check
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gate: must have at least one table with correct dimensions
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    if len(table.rows) < 6 or len(table.columns) < 3:
        print(f"CRITICAL: Table too small: {len(table.rows)} rows x {len(table.columns)} cols, expected 6x3")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify surrounding cells are intact
    # If this fails, the file is corrupted — return 0.0 immediately
    corrupted = []
    for (row_idx, col_idx), expected_val in PRECONDITION_CELLS.items():
        cell_text = table.cell(row_idx, col_idx).text.strip()
        if cell_text != expected_val:
            corrupted.append(f"[{row_idx},{col_idx}] expected '{expected_val}', found '{cell_text}'")
    if corrupted:
        print("CRITICAL: Precondition gate failed — surrounding cells are corrupted:")
        for c in corrupted:
            print(f"  {c}")
        print("REWARD: 0.0")
        return 0.0
    print("PRECONDITION: Surrounding cells intact — proceeding with scoring")

    # --- Component 1: Cell [5,2] displays the value '1025' (0.5 points) ---
    # This FAILS on initial_env (empty cell) and PASSES on golden_env (1025)
    try:
        last_cell = table.cell(5, 2)
        cell_text = last_cell.text.strip()
        if cell_text == EXPECTED_TOTAL:
            print(f"PASS: Component 1 — Cell [5,2] displays '{EXPECTED_TOTAL}' (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Cell [5,2] expected '{EXPECTED_TOTAL}', found '{repr(cell_text)}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: Cell [5,2] contains a Writer field formula using SUM (0.5 points) ---
    # This FAILS on initial_env (empty cell, no field) and PASSES on golden_env (=SUM(ABOVE))
    try:
        last_cell = table.cell(5, 2)
        cell_xml = last_cell._element

        # Look for w:fldChar elements — these indicate a Writer field formula
        fld_chars = cell_xml.findall(f'.//{{{W_NS}}}fldChar')
        # Look for w:instrText — contains the formula instruction text
        instr_texts = cell_xml.findall(f'.//{{{W_NS}}}instrText')

        has_field = len(fld_chars) >= 2  # at least begin + end field characters
        formula_text = ' '.join(
            (el.text or '').strip() for el in instr_texts
        ).upper()
        has_sum_formula = 'SUM' in formula_text

        if has_field and has_sum_formula:
            print(f"PASS: Component 2 — Writer SUM formula found: '{formula_text.strip()}' (0.5 pts)")
            total_score += 0.5
        elif has_field and not has_sum_formula:
            print(f"FAIL: Component 2 — Field present but not a SUM formula: '{formula_text.strip()}'")
        else:
            print(f"FAIL: Component 2 — No Writer field formula (w:fldChar + w:instrText) found in cell [5,2]")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


file_path = os.path.join(WORKDIR, TASK_FILE)
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
