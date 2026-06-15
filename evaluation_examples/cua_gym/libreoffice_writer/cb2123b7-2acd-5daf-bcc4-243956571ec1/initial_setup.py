"""
Initial Setup: Standard Business Report - Wrong Page Settings
Task ID: writer_page_049
Domain: libreoffice_writer

Creates a standard_report.docx on ~/Desktop/ with:
  - Letter size paper (8.5" x 11")
  - Landscape orientation
  - Margins: top=1.91cm, bottom=1.91cm, left=1.91cm, right=1.91cm
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'writer_page_049'
FILENAME = 'standard_report.docx'
OUTPUT = f'{DESKTOP}/{FILENAME}'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # -----------------------------------------------------------------------
    # Page setup: Letter size, Landscape, margins 1.91cm all sides
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    # Letter size: 8.5" x 11"
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Landscape: swap width/height and set orientation
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    # Margins 1.91cm on all sides
    section.top_margin = Cm(1.91)
    section.bottom_margin = Cm(1.91)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    # -----------------------------------------------------------------------
    # Cover / Title Section
    # -----------------------------------------------------------------------
    title = doc.add_heading('Annual Operations Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph('Meridian Solutions Group')
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)

    date_para = doc.add_paragraph('Fiscal Year 2024 | Q4 Summary')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # spacer

    # -----------------------------------------------------------------------
    # Executive Summary (Page 1 cont.)
    # -----------------------------------------------------------------------
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'Meridian Solutions Group achieved record-breaking performance in fiscal year 2024, '
        'surpassing projected targets across all major business units. This report provides '
        'a comprehensive overview of financial results, operational milestones, strategic '
        'initiatives, and forward-looking projections for the upcoming fiscal cycle.'
    )
    doc.add_paragraph(
        'Revenue grew by 18.4% year-over-year, driven primarily by expansion in the Asia-Pacific '
        'region and the successful launch of three new product lines. Operating costs were reduced '
        'by 6.2% through process optimization programs initiated in Q2 2024, resulting in an '
        'overall EBITDA margin improvement of 4.1 percentage points.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Financial Overview (Page 2)
    # -----------------------------------------------------------------------
    doc.add_heading('2. Financial Overview', level=1)
    doc.add_paragraph(
        'The following section details the consolidated financial performance of Meridian Solutions '
        'Group for the period January 1 – December 31, 2024.'
    )

    doc.add_heading('2.1 Revenue by Business Unit', level=2)
    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    headers = ['Business Unit', 'FY2023 Revenue ($M)', 'FY2024 Revenue ($M)']
    for i, h in enumerate(headers):
        run = table1.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    data = [
        ['North America Operations', '142.8', '168.3'],
        ['EMEA Division', '98.4', '115.7'],
        ['Asia-Pacific Expansion', '54.1', '78.9'],
        ['Digital Services', '37.6', '51.2'],
        ['Corporate & Shared Services', '-12.3', '-14.1'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table1.cell(r, c).text = val

    doc.add_paragraph()

    doc.add_heading('2.2 Cost Structure Analysis', level=2)
    doc.add_paragraph(
        'Total operating expenses for FY2024 amounted to $287.4 million, compared to $285.1 million '
        'in FY2023. Despite revenue growth, absolute cost levels remained largely stable due to '
        'efficiency programs across procurement, logistics, and corporate overhead.'
    )
    doc.add_paragraph(
        'Key cost reduction areas included: supplier renegotiation saving $8.3M, automation of '
        'back-office functions saving $4.7M, and rationalization of the real estate portfolio '
        'generating $2.1M in annual savings.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Operations Report (Page 3)
    # -----------------------------------------------------------------------
    doc.add_heading('3. Operational Performance', level=1)
    doc.add_paragraph(
        'Operational KPIs demonstrated consistent improvement across all measured dimensions. '
        'Customer satisfaction scores reached an all-time high of 87.3%, while average order '
        'fulfillment time decreased from 4.2 days to 3.1 days year-over-year.'
    )

    doc.add_heading('3.1 Key Performance Indicators', level=2)
    table2 = doc.add_table(rows=7, cols=4)
    table2.style = 'Table Grid'
    kpi_headers = ['KPI', 'Target', 'FY2024 Actual', 'Status']
    for i, h in enumerate(kpi_headers):
        run = table2.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    kpi_data = [
        ['Customer Satisfaction Score', '85%', '87.3%', 'Achieved'],
        ['On-Time Delivery Rate', '95%', '96.8%', 'Achieved'],
        ['Defect Rate (PPM)', '<500', '312', 'Achieved'],
        ['Employee Engagement Index', '75%', '79.1%', 'Achieved'],
        ['Net Promoter Score', '45', '52', 'Achieved'],
        ['Cost per Unit (avg)', '$12.40', '$11.87', 'Achieved'],
    ]
    for r, row_data in enumerate(kpi_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Human Resources (Page 4)
    # -----------------------------------------------------------------------
    doc.add_heading('4. Human Resources & Organizational Development', level=1)
    doc.add_paragraph(
        'As of December 31, 2024, Meridian Solutions Group employed 3,847 full-time equivalent '
        'staff globally, representing a net increase of 213 employees compared to year-end 2023. '
        'Voluntary attrition declined to 9.4%, the lowest level in five years.'
    )

    doc.add_heading('4.1 Workforce Composition', level=2)
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Table Grid'
    hr_headers = ['Department', 'Headcount', 'YoY Change']
    for i, h in enumerate(hr_headers):
        run = table3.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    hr_data = [
        ['Engineering & Technology', '1,142', '+87'],
        ['Sales & Marketing', '689', '+34'],
        ['Operations & Logistics', '1,215', '+61'],
        ['Finance & Accounting', '243', '+12'],
        ['Human Resources & Legal', '558', '+19'],
    ]
    for r, row_data in enumerate(hr_data, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'Training and development investment increased to $4.2 million in FY2024, up from '
        '$3.6 million in FY2023. New programs introduced include a leadership acceleration '
        'track for high-potential managers, technical certification pathways for engineering '
        'staff, and a company-wide digital literacy initiative.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Strategic Initiatives (Page 5)
    # -----------------------------------------------------------------------
    doc.add_heading('5. Strategic Initiatives', level=1)
    doc.add_paragraph(
        'In alignment with the Board-approved five-year strategy, Meridian Solutions Group '
        'executed several transformational initiatives during FY2024 that position the company '
        'for sustained long-term growth.'
    )

    doc.add_heading('5.1 Digital Transformation Program', level=2)
    doc.add_paragraph(
        'Phase 2 of the Digital Transformation Program was completed in September 2024, '
        'delivering a fully integrated ERP system across all North American and EMEA operations. '
        'The rollout covered 2,100 users and replaced 14 legacy systems, reducing IT operational '
        'costs by an estimated $3.8 million annually.'
    )

    doc.add_heading('5.2 Sustainability & ESG Commitment', level=2)
    doc.add_paragraph(
        'Meridian achieved a 22% reduction in Scope 1 and Scope 2 carbon emissions compared '
        'to the 2021 baseline year, ahead of the 2025 target of 20%. Renewable energy now '
        'constitutes 41% of total energy consumption across all company-owned facilities.'
    )
    doc.add_paragraph(
        'The company was awarded a BBB ESG rating by the Sustainable Finance Institute, '
        'representing a two-notch upgrade from the prior year assessment.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Risk Management (Page 6)
    # -----------------------------------------------------------------------
    doc.add_heading('6. Risk Management', level=1)
    doc.add_paragraph(
        'The Risk Management Committee conducted four formal reviews during FY2024, identifying '
        'and assessing 38 material risk items. Of these, 12 were classified as high-priority and '
        'assigned to dedicated cross-functional mitigation teams.'
    )

    doc.add_heading('6.1 Top Enterprise Risks', level=2)
    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Table Grid'
    risk_headers = ['Risk Category', 'Likelihood', 'Impact', 'Mitigation Status']
    for i, h in enumerate(risk_headers):
        run = table4.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    risk_data = [
        ['Cybersecurity Breach', 'Medium', 'High', 'In Progress'],
        ['Supply Chain Disruption', 'Medium', 'High', 'Completed'],
        ['Regulatory Non-Compliance', 'Low', 'High', 'Completed'],
        ['Key Talent Retention', 'Medium', 'Medium', 'In Progress'],
        ['Geopolitical Instability', 'High', 'Medium', 'Monitoring'],
    ]
    for r, row_data in enumerate(risk_data, 1):
        for c, val in enumerate(row_data):
            table4.cell(r, c).text = val

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Technology & Innovation (Page 7)
    # -----------------------------------------------------------------------
    doc.add_heading('7. Technology & Innovation', level=1)
    doc.add_paragraph(
        'Research and Development expenditure reached $18.7 million in FY2024, representing '
        '4.7% of total revenue. Three product innovations were commercially launched, with '
        'two additional patents filed in advanced manufacturing processes.'
    )

    doc.add_heading('7.1 Product Pipeline', level=2)
    doc.add_paragraph(
        'The product development pipeline contains 11 active projects across three maturity stages: '
        'early-stage concept (4 projects), development and validation (5 projects), and '
        'pre-commercial launch (2 projects). Combined addressable market value is estimated at $340 million.'
    )
    doc.add_paragraph(
        'Notable highlights include the MeriFlow 3.0 logistics optimization platform, scheduled '
        'for commercial launch in Q2 2025, and an AI-driven quality inspection system currently '
        'in final validation trials across two manufacturing facilities.'
    )

    doc.add_heading('7.2 Technology Infrastructure', level=2)
    doc.add_paragraph(
        'Capital expenditure on technology infrastructure totaled $22.4 million, primarily '
        'directed toward cloud migration (AWS and Azure hybrid model), cybersecurity enhancement, '
        'and data center consolidation from six facilities to three regional hubs.'
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Outlook (Page 8)
    # -----------------------------------------------------------------------
    doc.add_heading('8. FY2025 Outlook & Forward Guidance', level=1)
    doc.add_paragraph(
        'Management provides the following preliminary financial guidance for fiscal year 2025, '
        'subject to macroeconomic conditions and the successful execution of identified growth initiatives.'
    )

    doc.add_heading('8.1 Financial Targets', level=2)
    table5 = doc.add_table(rows=5, cols=3)
    table5.style = 'Table Grid'
    outlook_headers = ['Financial Metric', 'FY2024 Actual', 'FY2025 Guidance']
    for i, h in enumerate(outlook_headers):
        run = table5.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    outlook_data = [
        ['Total Revenue', '$399.0M', '$430M – $445M'],
        ['Gross Margin', '41.3%', '42.0% – 43.5%'],
        ['EBITDA', '$87.2M', '$95M – $102M'],
        ['Capital Expenditure', '$31.6M', '$28M – $33M'],
    ]
    for r, row_data in enumerate(outlook_data, 1):
        for c, val in enumerate(row_data):
            table5.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_heading('8.2 Strategic Priorities for FY2025', level=2)
    doc.add_paragraph(
        'Our strategic agenda for FY2025 is anchored on four priorities: (1) Accelerating Asia-Pacific '
        'market penetration through two new distribution partnerships, (2) Completing Phase 3 of the '
        'Digital Transformation Program covering the Asia-Pacific ERP rollout, (3) Launching the '
        'MeriFlow 3.0 platform and expanding the Digital Services client base by 30%, and '
        '(4) Achieving net-zero Scope 1 emissions from company-owned transport fleet.'
    )

    doc.add_paragraph()
    closing = doc.add_paragraph(
        'This report has been reviewed and approved by the Executive Leadership Team and the '
        'Board of Directors. Figures are unaudited and subject to final year-end audit adjustments.'
    )
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
