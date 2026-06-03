"""
Initial Setup: Book review document for indent task
Task ID: writer_para_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm

WORKDIR = '/home/user'
TASK_ID = 'writer_para_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1
    heading = doc.add_heading('Book Review: The Invisible Hand of Data', level=1)

    # Paragraph 2: Normal paragraph (no indent)
    p2 = doc.add_paragraph(
        'In her latest work, Professor Sarah Chen argues that big data analytics have '
        'fundamentally altered the balance of power between corporations and consumers:'
    )

    # Paragraph 3: Excerpt paragraph — NO indent in initial state (task is to add 3cm left, 1cm right)
    p3 = doc.add_paragraph(
        'We have entered an era where every click, every purchase, every moment of hesitation '
        'before a decision is captured, analyzed, and monetized. The consumer, once the sovereign '
        'of the marketplace, has become its most valuable product.'
    )
    # Explicitly ensure no indent on paragraph 3
    p3.paragraph_format.left_indent = Cm(0)
    p3.paragraph_format.right_indent = Cm(0)

    # Paragraph 4: Normal paragraph (no indent)
    p4 = doc.add_paragraph(
        'This provocative thesis is supported by extensive case studies from the technology sector, '
        'including a detailed analysis of targeted advertising algorithms:'
    )

    # Paragraph 5: Normal paragraph (no indent)
    p5 = doc.add_paragraph(
        'Chen presents compelling evidence that resonates with everyday experience:'
    )

    # Paragraph 6: Excerpt paragraph — NO indent in initial state (task is to add 3cm left, 1cm right)
    p6 = doc.add_paragraph(
        'The algorithms do not merely predict what we want \u2014 they shape what we want. '
        'The distinction between discovery and manipulation has become vanishingly thin, and we '
        'lack both the tools and the will to see the difference.'
    )
    # Explicitly ensure no indent on paragraph 6
    p6.paragraph_format.left_indent = Cm(0)
    p6.paragraph_format.right_indent = Cm(0)

    # Paragraph 7: Normal paragraph (no indent)
    p7 = doc.add_paragraph(
        'Overall, this is an essential read for anyone concerned about digital privacy and '
        'consumer rights in the modern economy.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
