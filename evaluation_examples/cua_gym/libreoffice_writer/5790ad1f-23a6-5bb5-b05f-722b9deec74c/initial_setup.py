"""
Initial Setup: Disciplinary warning letter with 'FINAL WARNING' in default formatting
Task ID: writer_hr_011
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_011'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # --- "FINAL WARNING" heading - default black, NOT bold ---
    warning_para = doc.add_paragraph()
    warning_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    warning_para.paragraph_format.space_after = Pt(12)
    run = warning_para.add_run("FINAL WARNING")
    run.font.size = Pt(16)
    run.font.name = 'Calibri'
    run.bold = False
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black - NOT red

    # --- Letter header ---
    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_after = Pt(6)
    header_para.add_run("Meridian Technologies, Inc.").bold = True
    header_para.add_run("\nHuman Resources Department")
    header_para.add_run("\n4200 Innovation Boulevard, Suite 300")
    header_para.add_run("\nAustin, TX 78759")

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(12)
    date_para.paragraph_format.space_after = Pt(12)
    date_para.add_run("Date: March 18, 2026")

    # --- Employee info ---
    emp_para = doc.add_paragraph()
    emp_para.paragraph_format.space_after = Pt(6)
    emp_para.add_run("To: ").bold = True
    emp_para.add_run("Derek Lawson")
    emp_para.add_run("\nEmployee ID: ").bold = True
    emp_para.add_run("MT-20234871")
    emp_para.add_run("\nDepartment: ").bold = True
    emp_para.add_run("Customer Support")
    emp_para.add_run("\nPosition: ").bold = True
    emp_para.add_run("Senior Support Specialist")

    # --- Subject ---
    subject_para = doc.add_paragraph()
    subject_para.paragraph_format.space_before = Pt(12)
    subject_para.paragraph_format.space_after = Pt(6)
    run_subj = subject_para.add_run("Subject: Final Written Warning - Repeated Policy Violations")
    run_subj.bold = True
    run_subj.underline = True

    # --- Body paragraph 1 ---
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(6)
    body1.add_run(
        "Dear Mr. Lawson,"
    )

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(6)
    body2.add_run(
        "This letter serves as your final written warning regarding ongoing violations of company "
        "policies, specifically the Attendance and Punctuality Policy (Section 4.2) and the "
        "Workplace Conduct Standards (Section 7.1) of the Meridian Technologies Employee Handbook. "
        "This warning follows two previous verbal counseling sessions held on January 10, 2026, "
        "and February 5, 2026, as well as a written warning issued on February 19, 2026."
    )

    # --- Body paragraph 2: violations ---
    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(6)
    body3.add_run("Specific Incidents:").bold = True

    violations = [
        "January 8, 2026 - Unexcused absence without prior notification to supervisor.",
        "January 22, 2026 - Late arrival (45 minutes) to scheduled shift; no call-in recorded.",
        "February 12, 2026 - Failure to follow escalation procedures on three customer tickets (IDs: #78432, #78455, #78491), resulting in SLA breaches.",
        "March 3, 2026 - Inappropriate language used during a team meeting, witnessed by two colleagues.",
        "March 14, 2026 - Unauthorized early departure (left 2 hours before end of shift) without manager approval.",
    ]
    for v in violations:
        doc.add_paragraph(v, style='List Bullet')

    # --- Expectations ---
    body4 = doc.add_paragraph()
    body4.paragraph_format.space_before = Pt(6)
    body4.paragraph_format.space_after = Pt(6)
    body4.add_run(
        "You are hereby placed on a 90-day performance improvement period effective immediately. "
        "During this period, you are expected to maintain full compliance with all company policies, "
        "arrive on time for every scheduled shift, follow all departmental procedures without exception, "
        "and demonstrate professional conduct at all times."
    )

    # --- Consequences ---
    body5 = doc.add_paragraph()
    body5.paragraph_format.space_after = Pt(6)
    body5.add_run(
        "Please be advised that any further policy violations during this improvement period will "
        "result in immediate termination of your employment with Meridian Technologies, Inc. This "
        "decision is final and is not subject to additional progressive disciplinary steps."
    )

    # --- Acknowledgment ---
    body6 = doc.add_paragraph()
    body6.paragraph_format.space_before = Pt(12)
    body6.paragraph_format.space_after = Pt(6)
    body6.add_run(
        "Please sign below to acknowledge that you have received and understood this final warning. "
        "Your signature does not indicate agreement with the above statements, only that you have "
        "been informed of the contents of this letter."
    )

    # --- Signature lines ---
    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_before = Pt(24)
    sig_para.add_run("_______________________________\n")
    sig_para.add_run("Employee Signature\t\t\tDate: _______________")

    sig_para2 = doc.add_paragraph()
    sig_para2.paragraph_format.space_before = Pt(18)
    sig_para2.add_run("_______________________________\n")
    sig_para2.add_run("Patricia Mendes\nHR Director, Meridian Technologies, Inc.")

    sig_para3 = doc.add_paragraph()
    sig_para3.paragraph_format.space_before = Pt(18)
    sig_para3.add_run("_______________________________\n")
    sig_para3.add_run("James Okonkwo\nCustomer Support Manager")

    # --- CC ---
    cc_para = doc.add_paragraph()
    cc_para.paragraph_format.space_before = Pt(12)
    cc_para.add_run("CC: ").bold = True
    cc_para.add_run("Employee Personnel File, Legal Department")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
