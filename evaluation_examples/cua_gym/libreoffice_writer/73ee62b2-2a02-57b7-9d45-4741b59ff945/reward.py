"""
Reward Script: Remove bullet formatting from items 3 and 4 in bullet list
Task ID: writer_list_070
Domain: libreoffice_writer
Scoring:
  Component 1: Items 3 & 4 ("Note:" paragraphs) are no longer List Bullet style (0.50 pts)
  Component 2: Items 3 & 4 have left indentation (matching bullet text level) (0.25 pts)
  Component 3: Items 1, 2, 5, 6 (index 0,1,4,5) remain as List Bullet style (0.25 pts)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_list_070'
FILE_PATH = f'{WORKDIR}/Desktop/mixed_content.docx'


def is_list_bullet(para):
    """Return True if the paragraph uses List Bullet style (via pStyle or style.name)."""
    # Check style name via python-docx
    if para.style and 'List Bullet' in para.style.name:
        return True
    # Also check direct pStyle element in pPr XML
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            val = pStyle.get(qn('w:val'), '')
            if 'ListBullet' in val or 'List Bullet' in val:
                return True
    return False


def has_left_indent(para):
    """Return True if the paragraph has a meaningful left indent (>0 EMU)."""
    pf = para.paragraph_format
    left_indent = pf.left_indent
    if left_indent is not None and left_indent > 0:
        return True
    # Also check XML directly for w:ind w:left
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            left_val = ind.get(qn('w:left'))
            if left_val is not None:
                try:
                    return int(left_val) > 0
                except ValueError:
                    pass
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify file has exactly 6 paragraphs
    paras = [p for p in doc.paragraphs if p.text.strip()]
    if len(paras) < 6:
        print(f"PRECONDITION FAIL: Expected 6 content paragraphs, found {len(paras)}")
        print(f"REWARD: 0.0")
        return 0.0

    # Use doc.paragraphs directly (all 6 are content paragraphs)
    all_paras = doc.paragraphs
    if len(all_paras) < 6:
        print(f"PRECONDITION FAIL: Document has fewer than 6 paragraphs: {len(all_paras)}")
        print("REWARD: 0.0")
        return 0.0

    para_2 = all_paras[2]  # "Note: Timeline subject to change..."
    para_3 = all_paras[3]  # "Note: Budget allocation..."

    # Verify the texts are what we expect (just for logging)
    print(f"Para 2 text: {para_2.text!r}")
    print(f"Para 3 text: {para_3.text!r}")

    # Component 1: Items 3 & 4 (index 2 and 3) are no longer List Bullet style (0.50 points)
    # In initial_env: both have style='List Bullet'
    # In golden_env: both should be 'Normal' or any non-bullet style
    try:
        para2_not_bullet = not is_list_bullet(para_2)
        para3_not_bullet = not is_list_bullet(para_3)

        if para2_not_bullet and para3_not_bullet:
            print(f"PASS: Component 1 — Items 3 & 4 have bullet formatting removed "
                  f"(styles: {para_2.style.name!r}, {para_3.style.name!r}) (0.50 pts)")
            total_score += 0.50
        elif para2_not_bullet:
            print(f"PARTIAL: Component 1 — Only item 3 has bullet removed "
                  f"(style: {para_2.style.name!r}), item 4 still has bullet (style: {para_3.style.name!r}) (0.25 pts)")
            total_score += 0.25
        elif para3_not_bullet:
            print(f"PARTIAL: Component 1 — Only item 4 has bullet removed "
                  f"(style: {para_3.style.name!r}), item 3 still has bullet (style: {para_2.style.name!r}) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Items 3 & 4 still have bullet style "
                  f"(styles: {para_2.style.name!r}, {para_3.style.name!r})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Items 3 & 4 have left indentation (0.25 points)
    # Task requires matching indentation of bullet text level to maintain visual alignment
    try:
        para2_indented = has_left_indent(para_2)
        para3_indented = has_left_indent(para_3)

        if para2_indented and para3_indented:
            indent2 = para_2.paragraph_format.left_indent
            indent3 = para_3.paragraph_format.left_indent
            print(f"PASS: Component 2 — Items 3 & 4 have left indentation "
                  f"(indent values: {indent2}, {indent3}) (0.25 pts)")
            total_score += 0.25
        elif para2_indented or para3_indented:
            print(f"PARTIAL: Component 2 — Only one of items 3 & 4 has indentation "
                  f"(para2_indented={para2_indented}, para3_indented={para3_indented}) (0.12 pts)")
            total_score += 0.12
        else:
            print(f"FAIL: Component 2 — Items 3 & 4 lack left indentation "
                  f"(para2_indent={para_2.paragraph_format.left_indent}, "
                  f"para3_indent={para_3.paragraph_format.left_indent})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Items 1, 2, 5, 6 (index 0,1,4,5) remain as List Bullet AND items 3&4 are NOT bullets (0.25 points)
    # This is a compound check: the bullet removal was selective (only items 3&4 changed, not others).
    # In initial_env: items 0,1,4,5 are List Bullet BUT items 2,3 are ALSO List Bullet → compound FAILS
    # In golden_env: items 0,1,4,5 are List Bullet AND items 2,3 are NOT List Bullet → compound PASSES
    try:
        bullet_indices = [0, 1, 4, 5]
        bullet_paras = [all_paras[i] for i in bullet_indices]
        all_still_bullet = all(is_list_bullet(p) for p in bullet_paras)
        still_bullet_count = sum(1 for p in bullet_paras if is_list_bullet(p))
        # Also confirm items 2 and 3 are NOT bullets (ensuring selective removal)
        notes_are_not_bullets = (not is_list_bullet(para_2)) and (not is_list_bullet(para_3))

        if all_still_bullet and notes_are_not_bullets:
            print(f"PASS: Component 3 — Items 0,1,4,5 retain List Bullet style AND "
                  f"items 2,3 are confirmed non-bullet (selective removal verified) (0.25 pts)")
            total_score += 0.25
        elif all_still_bullet and not notes_are_not_bullets:
            print(f"FAIL: Component 3 — Items 0,1,4,5 retain List Bullet BUT items 2 or 3 still have bullet "
                  f"(styles: {para_2.style.name!r}, {para_3.style.name!r})")
        elif notes_are_not_bullets and still_bullet_count >= 3:
            print(f"PARTIAL: Component 3 — Only {still_bullet_count}/4 bullet items retained List Bullet style (0.12 pts)")
            total_score += 0.12
        else:
            styles = [p.style.name for p in bullet_paras]
            print(f"FAIL: Component 3 — {still_bullet_count}/4 bullet items retain List Bullet style "
                  f"and/or items 2,3 not confirmed as non-bullet (styles: {styles})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
