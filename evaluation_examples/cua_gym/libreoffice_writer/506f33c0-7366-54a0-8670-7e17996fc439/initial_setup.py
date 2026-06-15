"""
Initial Setup: Format project phases as uppercase-letter numbered list
Task ID: writer_lec_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Project Plan', level=0)

    # Introduction paragraph
    doc.add_paragraph(
        'This document outlines the key phases of the Horizon Analytics Platform '
        'migration project. Each phase includes defined milestones, resource '
        'allocations, and success criteria that must be met before proceeding '
        'to the next stage.'
    )

    # Project phases as plain paragraphs (NO list formatting)
    doc.add_paragraph('Planning')
    doc.add_paragraph('Development')
    doc.add_paragraph('Testing')
    doc.add_paragraph('Deployment')

    # Additional context after phases
    doc.add_paragraph(
        'The estimated timeline for all phases is 18 months, with quarterly '
        'reviews scheduled to assess progress and adjust resource allocation '
        'as needed. Budget approval for each phase is contingent on successful '
        'completion of the preceding phase.'
    )

    doc.add_paragraph(
        'For questions regarding the project plan, contact the Program '
        'Management Office at pmo@horizonanalytics.com.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
