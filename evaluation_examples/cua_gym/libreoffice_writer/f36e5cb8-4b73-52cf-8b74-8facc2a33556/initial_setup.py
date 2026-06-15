"""
Initial Setup: Apply 'keep with next' to figure caption paragraphs
Task ID: writer_para_049
Domain: libreoffice_writer

Creates science_report.docx with 8 paragraphs matching the task context.
Figure captions (paragraphs 3, 5, 7) do NOT yet have keep_with_next set.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_para_049'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1
    heading = doc.add_heading('Experimental Results: Photovoltaic Cell Efficiency', level=1)
    # Ensure keep_with_next is not set (default/None)

    # Paragraph 2: Introductory paragraph
    para2 = doc.add_paragraph(
        'The following figures present our key experimental findings from the laboratory tests '
        'conducted between January and June 2024.'
    )

    # Paragraph 3: Figure 1 caption (italic) — keep_with_next NOT set
    para3 = doc.add_paragraph()
    run3 = para3.add_run('Figure 1: Efficiency vs. Temperature Curve')
    run3.italic = True

    # Paragraph 4: Figure 1 description
    para4 = doc.add_paragraph(
        'The efficiency of the perovskite solar cell showed a linear decrease of 0.45% per degree '
        'Celsius increase above 25\u00b0C, consistent with theoretical predictions.'
    )

    # Paragraph 5: Figure 2 caption (italic) — keep_with_next NOT set
    para5 = doc.add_paragraph()
    run5 = para5.add_run('Figure 2: Degradation Rate Over Time')
    run5.italic = True

    # Paragraph 6: Figure 2 description
    para6 = doc.add_paragraph(
        'Under continuous illumination at AM1.5 conditions, the cell retained 92% of its initial '
        'efficiency after 500 hours, significantly outperforming previous generation devices.'
    )

    # Paragraph 7: Figure 3 caption (italic) — keep_with_next NOT set
    para7 = doc.add_paragraph()
    run7 = para7.add_run('Figure 3: Spectral Response Comparison')
    run7.italic = True

    # Paragraph 8: Figure 3 description
    para8 = doc.add_paragraph(
        'The spectral response curve shows enhanced absorption in the 400-700nm range compared to '
        'silicon reference cells.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
