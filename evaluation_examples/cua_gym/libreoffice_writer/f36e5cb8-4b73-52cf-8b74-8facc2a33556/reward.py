"""
Reward Script: Apply 'keep with next' setting to figure caption paragraphs
Task ID: writer_para_049
Domain: libreoffice_writer
Scoring:
  Component 1: Figure 1 caption (Para 2) has keep_with_next=True (0.33 pts)
  Component 2: Figure 2 caption (Para 4) has keep_with_next=True (0.34 pts)
  Component 3: Figure 3 caption (Para 6) has keep_with_next=True (0.33 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_para_049'

# Expected figure captions by paragraph index and their partial text
FIGURE_CAPTIONS = {
    2: 'Figure 1: Efficiency vs. Temperature Curve',
    4: 'Figure 2: Degradation Rate Over Time',
    6: 'Figure 3: Spectral Response Comparison',
}


def has_keep_with_next(para):
    """
    Check if a paragraph has keep_with_next (keepNext) set to True.
    Checks via python-docx paragraph_format AND directly via XML to be thorough.
    """
    pf = para.paragraph_format
    if pf.keep_with_next is True:
        return True
    # Also check XML directly for robustness
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        kn = pPr.find(qn('w:keepNext'))
        if kn is not None:
            # keepNext element present means keep_with_next is effectively True
            # (unless val="0", check for that)
            val = kn.get(qn('w:val'))
            if val is None or val.lower() not in ('0', 'false'):
                return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Apply 'keep with next' to figure caption paragraphs (indices 2, 4, 6)
    so they always stay with the following description paragraph.
    """
    total_score = 0.0

    # Load document — precondition gate
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 8 paragraphs
    if len(doc.paragraphs) < 8:
        print(f"CRITICAL: Expected >= 8 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: verify text content is intact (no text changes allowed)
    paragraphs = doc.paragraphs
    expected_texts = {
        0: 'Experimental Results: Photovoltaic Cell Efficiency',
        1: 'The following figures present our key experimental findings from the laboratory tests conducted between January and June 2024.',
        2: 'Figure 1: Efficiency vs. Temperature Curve',
        3: 'The efficiency of the perovskite solar cell showed a linear decrease of 0.45% per degree Celsius increase above 25\u00b0C, consistent with theoretical predictions.',
        4: 'Figure 2: Degradation Rate Over Time',
        5: 'Under continuous illumination at AM1.5 conditions, the cell retained 92% of its initial efficiency after 500 hours, significantly outperforming previous generation devices.',
        6: 'Figure 3: Spectral Response Comparison',
        7: 'The spectral response curve shows enhanced absorption in the 400-700nm range compared to silicon reference cells.',
    }
    for idx, expected_text in expected_texts.items():
        actual_text = paragraphs[idx].text if idx < len(paragraphs) else ''
        if actual_text != expected_text:
            print(f"CRITICAL: Text content changed at paragraph {idx}. Expected: '{expected_text[:60]}', found: '{actual_text[:60]}'")
            print("REWARD: 0.0")
            return 0.0
    print("PASS: Text content unchanged (precondition met)")

    # Component 1: Figure 1 caption (Para 2) has keep_with_next=True (0.33 points)
    try:
        para2 = paragraphs[2]
        caption2_text = para2.text
        if not caption2_text.startswith('Figure 1'):
            print(f"FAIL: Component 1 — Paragraph 2 does not contain Figure 1 caption, found: '{caption2_text[:60]}'")
        elif has_keep_with_next(para2):
            print(f"PASS: Component 1 — Figure 1 caption has keep_with_next=True (0.33 pts)")
            total_score += 0.33
        else:
            kwn_val = para2.paragraph_format.keep_with_next
            print(f"FAIL: Component 1 — Figure 1 caption (Para 2) keep_with_next={kwn_val}, expected True")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Figure 2 caption (Para 4) has keep_with_next=True (0.34 points)
    try:
        para4 = paragraphs[4]
        caption4_text = para4.text
        if not caption4_text.startswith('Figure 2'):
            print(f"FAIL: Component 2 — Paragraph 4 does not contain Figure 2 caption, found: '{caption4_text[:60]}'")
        elif has_keep_with_next(para4):
            print(f"PASS: Component 2 — Figure 2 caption has keep_with_next=True (0.34 pts)")
            total_score += 0.34
        else:
            kwn_val = para4.paragraph_format.keep_with_next
            print(f"FAIL: Component 2 — Figure 2 caption (Para 4) keep_with_next={kwn_val}, expected True")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Figure 3 caption (Para 6) has keep_with_next=True (0.33 points)
    try:
        para6 = paragraphs[6]
        caption6_text = para6.text
        if not caption6_text.startswith('Figure 3'):
            print(f"FAIL: Component 3 — Paragraph 6 does not contain Figure 3 caption, found: '{caption6_text[:60]}'")
        elif has_keep_with_next(para6):
            print(f"PASS: Component 3 — Figure 3 caption has keep_with_next=True (0.33 pts)")
            total_score += 0.33
        else:
            kwn_val = para6.paragraph_format.keep_with_next
            print(f"FAIL: Component 3 — Figure 3 caption (Para 6) keep_with_next={kwn_val}, expected True")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Bonus integrity check: non-caption paragraphs must NOT have keep_with_next=True
    # (ensures only target paragraphs were modified, not over-application)
    non_caption_indices = [0, 1, 3, 5, 7]
    over_applied = []
    for idx in non_caption_indices:
        if idx < len(paragraphs) and has_keep_with_next(paragraphs[idx]):
            over_applied.append(idx)
    if over_applied:
        print(f"NOTE: Non-caption paragraphs {over_applied} also have keep_with_next=True (over-application, but not penalized)")
    else:
        print("PASS: Non-caption paragraphs correctly do not have keep_with_next set")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
