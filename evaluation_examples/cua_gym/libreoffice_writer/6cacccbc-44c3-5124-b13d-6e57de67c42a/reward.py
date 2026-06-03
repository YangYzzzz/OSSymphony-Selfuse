"""
Reward Script: FMLA Administration Guide
Task ID: writer_hr_075
Domain: libreoffice_writer
Scoring:
  C1: Multiple tables added (>= 6 tables)            — 0.15 pts
  C2: Heading structure (H1 and H2 styles used)       — 0.15 pts
  C3: Eligibility decision table (step-based)          — 0.15 pts
  C4: Leave calculation comparison table               — 0.10 pts
  C5: Three notice template sub-sections (3.1-3.3)     — 0.15 pts
  C6: Rights & responsibilities two-column table       — 0.10 pts
  C7: Certification form tables with fields            — 0.10 pts
  C8: Tracking log table (multi-column with data)      — 0.10 pts
  Total: 1.00
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_075'


def persist_app_state(domain: str):
    """Try to save any unsaved LibreOffice changes."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            import time
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_tables = len(doc.tables)
    num_paras = len(doc.paragraphs)

    # Collect heading info
    h1_texts = []
    h2_texts = []
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Heading 1':
            h1_texts.append(p.text.strip().lower())
        elif p.style and p.style.name == 'Heading 2':
            h2_texts.append(p.text.strip().lower())

    # ---------------------------------------------------------------
    # Component 1: Multiple tables added (>= 6 tables) — 0.15 pts
    # Initial has 0 tables; golden has 10.
    # ---------------------------------------------------------------
    try:
        if num_tables >= 6:
            print(f"PASS: Component 1 — {num_tables} tables found (>= 6) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — only {num_tables} tables found, need >= 6")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: Heading structure with H1 and H2 styles — 0.15 pts
    # Initial has all Normal style; golden has 7 H1 + 9 H2.
    # Award 0.08 for >= 5 H1, 0.07 for >= 3 H2.
    # ---------------------------------------------------------------
    try:
        c2_score = 0.0
        if len(h1_texts) >= 5:
            c2_score += 0.08
            print(f"  PASS: Component 2a — {len(h1_texts)} Heading 1 paragraphs (>= 5)")
        else:
            print(f"  FAIL: Component 2a — {len(h1_texts)} Heading 1 paragraphs, need >= 5")

        if len(h2_texts) >= 3:
            c2_score += 0.07
            print(f"  PASS: Component 2b — {len(h2_texts)} Heading 2 paragraphs (>= 3)")
        else:
            print(f"  FAIL: Component 2b — {len(h2_texts)} Heading 2 paragraphs, need >= 3")

        if c2_score > 0:
            print(f"PASS: Component 2 — heading structure ({c2_score} pts)")
            total_score += c2_score
        else:
            print(f"FAIL: Component 2 — no proper heading structure")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Eligibility decision table — 0.15 pts
    # Golden table 0 has columns: Step, Eligibility Criterion, Yes, No
    # At least 5 rows of step-based decision logic.
    # ---------------------------------------------------------------
    try:
        found_eligibility_table = False
        for table in doc.tables:
            if len(table.rows) < 5 or len(table.columns) < 3:
                continue
            header_text = ' '.join(cell.text.strip().lower() for cell in table.rows[0].cells)
            # Check for decision-tree like columns
            if ('step' in header_text or 'criterion' in header_text or 'eligib' in header_text):
                # Verify it has yes/no or proceed/result logic
                has_decision_logic = False
                for ri in range(1, min(len(table.rows), 5)):
                    row_text = ' '.join(cell.text.strip().lower() for cell in table.rows[ri].cells)
                    if ('go to' in row_text or 'not eligible' in row_text or
                            'eligible' in row_text or 'proceed' in row_text or
                            'step' in row_text):
                        has_decision_logic = True
                        break
                if has_decision_logic:
                    found_eligibility_table = True
                    break

        if found_eligibility_table:
            print(f"PASS: Component 3 — eligibility decision table found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — no eligibility decision table found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Leave calculation comparison table — 0.10 pts
    # Golden has a table comparing Rolling 12-Month, Calendar Year, Fixed Year methods.
    # ---------------------------------------------------------------
    try:
        found_calc_table = False
        for table in doc.tables:
            if len(table.rows) < 3 or len(table.columns) < 3:
                continue
            all_text = ' '.join(
                cell.text.strip().lower()
                for row in table.rows
                for cell in row.cells
            )
            # Check for method comparison keywords
            if ('rolling' in all_text and 'calendar' in all_text and
                    ('fixed' in all_text or 'method' in all_text)):
                found_calc_table = True
                break

        if found_calc_table:
            print(f"PASS: Component 4 — leave calculation comparison table found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — no leave calculation comparison table")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---------------------------------------------------------------
    # Component 5: Three notice template sub-sections — 0.15 pts
    # Golden has Heading 2 entries for 3.1, 3.2, 3.3 notice templates.
    # MUST be structured as headings (H2) or as distinct formatted sections
    # with template content (Date/To/From fields). Plain body mentions don't count.
    # ---------------------------------------------------------------
    try:
        notice_keywords = ['eligibility', 'designation', 'return']
        found_notices = set()

        # Primary: check H2 headings for notice sub-sections
        for h2 in h2_texts:
            for kw in notice_keywords:
                if kw in h2:
                    found_notices.add(kw)

        # Secondary: check H1 headings that might contain notice keywords
        for h1 in h1_texts:
            for kw in notice_keywords:
                if kw in h1 and 'notice' in h1:
                    found_notices.add(kw)

        # Tertiary: look for actual template structure (Date/To/From fields)
        # Only count if we find structured template content, not just mentions
        if len(found_notices) < 3:
            template_markers = 0
            for p in doc.paragraphs:
                text_lower = p.text.strip().lower()
                if (text_lower.startswith('date:') or text_lower.startswith('to:') or
                        text_lower.startswith('from:') or text_lower.startswith('re:')):
                    template_markers += 1
            # Need at least 6 template markers (2 per notice: Date+To+From per template)
            if template_markers >= 6:
                # Check which notices have template content by scanning nearby paragraphs
                in_notice_section = None
                for p in doc.paragraphs:
                    text_lower = p.text.strip().lower()
                    if 'notice of eligibility' in text_lower or 'eligibility and rights' in text_lower:
                        in_notice_section = 'eligibility'
                    elif 'designation notice' in text_lower:
                        in_notice_section = 'designation'
                    elif 'return-to-work' in text_lower or 'return to work' in text_lower:
                        in_notice_section = 'return'
                    if in_notice_section and (text_lower.startswith('date:') or text_lower.startswith('to:')):
                        found_notices.add(in_notice_section)

        if len(found_notices) >= 3:
            print(f"PASS: Component 5 — 3 notice templates as structured sections: {found_notices} (0.15 pts)")
            total_score += 0.15
        elif len(found_notices) >= 2:
            print(f"PARTIAL: Component 5 — {len(found_notices)} notice templates found: {found_notices} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — only {len(found_notices)} notice template sections found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ---------------------------------------------------------------
    # Component 6: Rights & responsibilities two-column table — 0.10 pts
    # Golden has a 2-column table: Employer vs Employee rights.
    # ---------------------------------------------------------------
    try:
        found_rights_table = False
        for table in doc.tables:
            if len(table.columns) != 2 or len(table.rows) < 4:
                continue
            header_text = ' '.join(cell.text.strip().lower() for cell in table.rows[0].cells)
            if ('employer' in header_text and 'employee' in header_text and
                    ('right' in header_text or 'responsib' in header_text)):
                found_rights_table = True
                break

        if found_rights_table:
            print(f"PASS: Component 6 — rights/responsibilities two-column table found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — no rights/responsibilities two-column table")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # ---------------------------------------------------------------
    # Component 7: Certification form tables with fields — 0.10 pts
    # Golden has 3 certification tables (WH-380-E, WH-380-F, Fitness).
    # Each is a 2-column table with "Field" and "Description" headers.
    # Need at least 2 such tables.
    # ---------------------------------------------------------------
    try:
        cert_table_count = 0
        for table in doc.tables:
            if len(table.columns) != 2 or len(table.rows) < 4:
                continue
            header_text = ' '.join(cell.text.strip().lower() for cell in table.rows[0].cells)
            if 'field' in header_text and ('description' in header_text or 'instruction' in header_text):
                # Verify it has certification-like fields
                all_text = ' '.join(
                    cell.text.strip().lower()
                    for row in table.rows
                    for cell in row.cells
                )
                if ('employee' in all_text or 'certification' in all_text or
                        'condition' in all_text or 'duration' in all_text):
                    cert_table_count += 1

        if cert_table_count >= 2:
            print(f"PASS: Component 7 — {cert_table_count} certification form tables found (0.10 pts)")
            total_score += 0.10
        elif cert_table_count >= 1:
            print(f"PARTIAL: Component 7 — {cert_table_count} certification form table (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 — no certification form tables found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # ---------------------------------------------------------------
    # Component 8: Tracking log table (multi-column with data) — 0.10 pts
    # Golden has a tracking table with >= 7 columns including Employee Name,
    # Leave Start, Expected Return, etc. and at least 3 data rows.
    # ---------------------------------------------------------------
    try:
        found_tracking_table = False
        for table in doc.tables:
            if len(table.columns) < 5 or len(table.rows) < 4:
                continue
            header_text = ' '.join(cell.text.strip().lower() for cell in table.rows[0].cells)
            if (('employee' in header_text or 'name' in header_text) and
                    ('leave' in header_text or 'start' in header_text or 'return' in header_text) and
                    ('reason' in header_text or 'status' in header_text or 'tracking' in header_text or
                     'days' in header_text or 'dept' in header_text)):
                # Check that it has actual data rows (not just headers)
                data_rows = 0
                for ri in range(1, len(table.rows)):
                    row_text = ' '.join(cell.text.strip() for cell in table.rows[ri].cells)
                    if len(row_text) > 10:
                        data_rows += 1
                if data_rows >= 3:
                    found_tracking_table = True
                    break

        if found_tracking_table:
            print(f"PASS: Component 8 — tracking log table with data found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 8 — no tracking log table with sufficient data")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
