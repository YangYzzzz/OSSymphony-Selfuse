"""
Initial Setup: FMLA Administration Guide - unformatted policy text
Task ID: writer_hr_075
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_075'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title - plain heading
    doc.add_heading('FMLA_Administration_Guide', level=0)

    # --- Section 1: Overview (unformatted text) ---
    doc.add_paragraph(
        'The Family and Medical Leave Act (FMLA) of 1993 provides eligible employees with up to '
        '12 weeks of unpaid, job-protected leave per year. It also requires that group health '
        'benefits be maintained during the leave period. FMLA applies to all public agencies, '
        'all public and private elementary and secondary schools, and companies with 50 or more '
        'employees within a 75-mile radius.'
    )

    doc.add_paragraph(
        'This administration guide is intended to serve as a comprehensive resource for HR '
        'professionals responsible for managing FMLA leave requests, determining eligibility, '
        'calculating leave entitlements, and ensuring compliance with federal regulations. '
        'The guide covers the full lifecycle of an FMLA case from initial request through '
        'return to work.'
    )

    # --- Section 2: Eligibility text (no decision tree/table) ---
    doc.add_paragraph(
        'Eligibility Requirements'
    )
    doc.add_paragraph(
        'To be eligible for FMLA leave, an employee must have worked for the employer for at '
        'least 12 months. The 12 months do not need to be consecutive. Additionally, the '
        'employee must have worked at least 1,250 hours during the 12-month period immediately '
        'preceding the start of the leave. The employer must have at least 50 employees within '
        'a 75-mile radius of the employee\'s worksite.'
    )
    doc.add_paragraph(
        'Qualifying reasons for FMLA leave include the birth of a child and bonding with the '
        'newborn within one year of birth, placement of a child for adoption or foster care and '
        'bonding within one year, caring for a spouse, child, or parent with a serious health '
        'condition, a serious health condition that makes the employee unable to perform the '
        'essential functions of their job, and any qualifying exigency arising from the fact '
        'that a spouse, child, or parent is a military member on covered active duty.'
    )

    # --- Section 3: Leave calculation text (no tables) ---
    doc.add_paragraph(
        'Leave Calculation Methods'
    )
    doc.add_paragraph(
        'Employers must select one of four methods for calculating the 12-month FMLA leave '
        'entitlement period. The rolling 12-month period is measured backward from the date '
        'an employee uses FMLA leave. Under this method, each time an employee takes FMLA '
        'leave, the remaining leave entitlement is the balance of the 12 weeks not used during '
        'the immediately preceding 12 months.'
    )
    doc.add_paragraph(
        'The calendar year method uses January 1 through December 31 as the measurement period. '
        'An employee is entitled to 12 weeks of FMLA leave at the start of each calendar year. '
        'The fixed leave year method uses a consistent 12-month period chosen by the employer, '
        'such as a fiscal year or the anniversary of the employee\'s hire date. The 12-month '
        'period measured forward from the first date of leave starts the calculation from the '
        'first day the employee takes FMLA leave.'
    )

    # --- Section 4: Notice requirements text (no templates) ---
    doc.add_paragraph(
        'Notice Requirements'
    )
    doc.add_paragraph(
        'Employers have specific notice obligations under FMLA. The general notice must be '
        'posted in a conspicuous place and included in employee handbooks. When an employee '
        'requests leave or the employer acquires knowledge that leave may be FMLA-qualifying, '
        'the employer must provide an eligibility notice within five business days. This notice '
        'must inform the employee whether they are eligible for FMLA leave.'
    )
    doc.add_paragraph(
        'A rights and responsibilities notice must be provided at the same time as the '
        'eligibility notice. This notice must include details about the expectations and '
        'obligations of the employee during leave. A designation notice must be provided '
        'within five business days after the employer has enough information to determine '
        'whether the leave qualifies as FMLA leave. If the leave is not designated as FMLA, '
        'the employer must notify the employee of the reason.'
    )

    # --- Section 5: Rights and responsibilities text (no two-column format) ---
    doc.add_paragraph(
        'Rights and Responsibilities'
    )
    doc.add_paragraph(
        'Employer responsibilities include maintaining group health insurance coverage during '
        'FMLA leave on the same terms as if the employee continued to work, restoring the '
        'employee to the same or equivalent position upon return, providing required notices '
        'in a timely manner, keeping accurate records of FMLA leave usage, and refraining '
        'from interfering with or retaliating against employees exercising FMLA rights.'
    )
    doc.add_paragraph(
        'Employee responsibilities include providing 30 days advance notice when the need '
        'for leave is foreseeable, providing sufficient information for the employer to '
        'determine if leave qualifies as FMLA, complying with the employer\'s usual call-in '
        'procedures, providing medical certification within 15 calendar days of the request, '
        'and periodically reporting on status and intent to return to work during the leave period.'
    )

    # --- Section 6: Certification text (no form fields) ---
    doc.add_paragraph(
        'Certification Requirements'
    )
    doc.add_paragraph(
        'The employer may require that the need for leave be supported by a certification '
        'issued by the health care provider of the employee or the family member. The '
        'certification must include the date the serious health condition commenced, the '
        'probable duration of the condition, appropriate medical facts regarding the condition, '
        'and a statement that the employee is needed to care for the family member or is unable '
        'to perform job functions. For intermittent leave, the certification must also include '
        'the expected frequency and duration of episodes.'
    )
    doc.add_paragraph(
        'Recertification may be requested every 30 days in connection with an absence unless '
        'the condition lasts longer than 30 days, in which case recertification may not be '
        'requested until the minimum duration expires. The employer may also require a '
        'fitness-for-duty certification before the employee returns to work.'
    )

    # --- Section 7: Tracking text (no table) ---
    doc.add_paragraph(
        'Leave Tracking'
    )
    doc.add_paragraph(
        'HR departments must maintain accurate records of all FMLA leave taken by employees. '
        'Records should include the employee name, department, leave start date, expected return '
        'date, actual return date, total hours or days used, qualifying reason, certification '
        'status, and any notes relevant to the case. These records must be retained for at '
        'least three years and be available for inspection by the Department of Labor.'
    )

    # --- Section 8: Related forms text (no cross-references) ---
    doc.add_paragraph(
        'Related Forms and References'
    )
    doc.add_paragraph(
        'Administering FMLA leave requires several standardized forms and documents. These '
        'include the initial leave request form, medical certification form, eligibility '
        'determination notice, rights and responsibilities notice, designation notice, '
        'return to work certification, and leave tracking spreadsheet. Each form should '
        'reference the applicable section of this guide and related federal regulations.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
