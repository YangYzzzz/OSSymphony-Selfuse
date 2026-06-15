"""
Reward Script: Mail merge records 3-7 from data source into invoice template
Task ID: writer_mt_009
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Document has exactly 5 tables (one per merged record)
  Component 2 (0.15): Document has 4 page breaks separating 5 pages
  Component 3 (0.30): Company names match records 3-7 in correct order
  Component 4 (0.20): Invoice numbers match records 3-7 in correct order
  Component 5 (0.15): Amounts and due dates match records 3-7
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_009'

# Expected data for records 3 through 7 (1-indexed from CSV)
EXPECTED_COMPANIES = [
    'Cascade Digital LLC',
    'Durango Manufacturing',
    'Evergreen Logistics',
    'Falcon Analytics Group',
    'Granite Peak Ventures',
]

EXPECTED_INVOICE_NUMS = [
    'INV-2025-003',
    'INV-2025-004',
    'INV-2025-005',
    'INV-2025-006',
    'INV-2025-007',
]

EXPECTED_AMOUNTS = [
    '$6,340.00',
    '$2,190.75',
    '$8,425.00',
    '$3,560.25',
    '$5,780.00',
]

EXPECTED_DATES = [
    '2025-05-01',
    '2025-05-08',
    '2025-05-15',
    '2025-05-22',
    '2025-06-01',
]


def verify_task(file_path):
    """
    Verify that the merged document contains exactly records 3-7,
    with correct data and page breaks.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Document has exactly 5 tables (one per merged record) (0.20 points)
    try:
        num_tables = len(doc.tables)
        if num_tables == 5:
            print(f"PASS: Component 1 — Found exactly 5 tables (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected 5 tables, found {num_tables}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Document has 4 page breaks separating 5 pages (0.15 points)
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        page_break_count = 0
        for para in doc.paragraphs:
            # Check run-level page breaks
            for run in para.runs:
                for br in run.element.findall('.//w:br', ns):
                    br_type = br.attrib.get(f'{{{ns["w"]}}}type', '')
                    if br_type == 'page':
                        page_break_count += 1
            # Check paragraph-level page break before
            if para.paragraph_format.page_break_before:
                page_break_count += 1

        if page_break_count == 4:
            print(f"PASS: Component 2 — Found exactly 4 page breaks (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Expected 4 page breaks, found {page_break_count}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Company names match records 3-7 in correct order (0.30 points)
    try:
        # Extract company names from "Bill To:" sections
        found_companies = []
        paragraphs = doc.paragraphs
        for i, para in enumerate(paragraphs):
            if para.text.strip() == 'Bill To:' and i + 1 < len(paragraphs):
                company = paragraphs[i + 1].text.strip()
                if company:
                    found_companies.append(company)

        if found_companies == EXPECTED_COMPANIES:
            print(f"PASS: Component 3 — All 5 companies match records 3-7 in order (0.30 pts)")
            total_score += 0.30
        else:
            # Partial credit: count matching companies
            matches = sum(1 for a, b in zip(found_companies, EXPECTED_COMPANIES) if a == b)
            if matches > 0 and len(found_companies) == 5:
                partial = 0.30 * (matches / 5)
                print(f"PARTIAL: Component 3 — {matches}/5 companies match ({partial:.2f} pts)")
                print(f"  Expected: {EXPECTED_COMPANIES}")
                print(f"  Found:    {found_companies}")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Companies don't match records 3-7")
                print(f"  Expected: {EXPECTED_COMPANIES}")
                print(f"  Found:    {found_companies}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Invoice numbers match records 3-7 (0.20 points)
    try:
        found_invoice_nums = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and 'Invoice Number:' in cells[0]:
                    found_invoice_nums.append(cells[1])

        if found_invoice_nums == EXPECTED_INVOICE_NUMS:
            print(f"PASS: Component 4 — All 5 invoice numbers match records 3-7 (0.20 pts)")
            total_score += 0.20
        else:
            matches = sum(1 for a, b in zip(found_invoice_nums, EXPECTED_INVOICE_NUMS) if a == b)
            if matches > 0 and len(found_invoice_nums) == 5:
                partial = 0.20 * (matches / 5)
                print(f"PARTIAL: Component 4 — {matches}/5 invoice numbers match ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Invoice numbers don't match records 3-7")
                print(f"  Expected: {EXPECTED_INVOICE_NUMS}")
                print(f"  Found:    {found_invoice_nums}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Amounts and due dates match records 3-7 (0.15 points)
    try:
        found_amounts = []
        found_dates = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    if 'Amount Due:' in cells[0]:
                        found_amounts.append(cells[1])
                    elif 'Due Date:' in cells[0]:
                        found_dates.append(cells[1])

        amounts_match = found_amounts == EXPECTED_AMOUNTS
        dates_match = found_dates == EXPECTED_DATES

        if amounts_match and dates_match:
            print(f"PASS: Component 5 — All amounts and dates match records 3-7 (0.15 pts)")
            total_score += 0.15
        elif amounts_match or dates_match:
            partial = 0.075
            which = "amounts" if amounts_match else "dates"
            print(f"PARTIAL: Component 5 — {which} match but not both ({partial:.2f} pts)")
            print(f"  Amounts expected: {EXPECTED_AMOUNTS}")
            print(f"  Amounts found:    {found_amounts}")
            print(f"  Dates expected:   {EXPECTED_DATES}")
            print(f"  Dates found:      {found_dates}")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Amounts and dates don't match records 3-7")
            print(f"  Amounts expected: {EXPECTED_AMOUNTS}")
            print(f"  Amounts found:    {found_amounts}")
            print(f"  Dates expected:   {EXPECTED_DATES}")
            print(f"  Dates found:      {found_dates}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
