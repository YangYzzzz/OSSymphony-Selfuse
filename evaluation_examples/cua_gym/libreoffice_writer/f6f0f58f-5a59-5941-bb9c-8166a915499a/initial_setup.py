"""
Initial Setup: Mail merge invoice template with data source
Task ID: writer_mt_009
Domain: libreoffice_writer
"""

import os
import csv
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DATA_SOURCE = f'{WORKDIR}/Invoices.csv'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# --- Invoice data source (12 records) ---
INVOICE_DATA = [
    {"Company": "Apex Solutions Inc.", "InvoiceNum": "INV-2025-001", "Amount": "$4,250.00", "DueDate": "2025-04-15"},
    {"Company": "BrightPath Consulting", "InvoiceNum": "INV-2025-002", "Amount": "$1,875.50", "DueDate": "2025-04-22"},
    {"Company": "Cascade Digital LLC", "InvoiceNum": "INV-2025-003", "Amount": "$6,340.00", "DueDate": "2025-05-01"},
    {"Company": "Durango Manufacturing", "InvoiceNum": "INV-2025-004", "Amount": "$2,190.75", "DueDate": "2025-05-08"},
    {"Company": "Evergreen Logistics", "InvoiceNum": "INV-2025-005", "Amount": "$8,425.00", "DueDate": "2025-05-15"},
    {"Company": "Falcon Analytics Group", "InvoiceNum": "INV-2025-006", "Amount": "$3,560.25", "DueDate": "2025-05-22"},
    {"Company": "Granite Peak Ventures", "InvoiceNum": "INV-2025-007", "Amount": "$5,780.00", "DueDate": "2025-06-01"},
    {"Company": "Harbor Point Trading", "InvoiceNum": "INV-2025-008", "Amount": "$1,950.00", "DueDate": "2025-06-08"},
    {"Company": "Ironclad Security", "InvoiceNum": "INV-2025-009", "Amount": "$7,120.50", "DueDate": "2025-06-15"},
    {"Company": "Jade Mountain Imports", "InvoiceNum": "INV-2025-010", "Amount": "$4,680.00", "DueDate": "2025-06-22"},
    {"Company": "Keystone Partners", "InvoiceNum": "INV-2025-011", "Amount": "$2,345.75", "DueDate": "2025-07-01"},
    {"Company": "Lakeshore Properties", "InvoiceNum": "INV-2025-012", "Amount": "$9,150.00", "DueDate": "2025-07-08"},
]


def create_data_source():
    """Create CSV data source with 12 invoice records."""
    with open(DATA_SOURCE, 'w', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=["Company", "InvoiceNum", "Amount", "DueDate"])
        writer.writeheader()
        writer.writerows(INVOICE_DATA)
    print(f'Data source created: {DATA_SOURCE}')


def create_template():
    """Create invoice template with merge field placeholders."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Company header area
    heading = doc.add_heading('INVOICE', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sep_run = sep.add_run('_' * 60)
    sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    sep_run.font.size = Pt(8)

    # Blank line
    doc.add_paragraph()

    # Bill To section
    bill_to_heading = doc.add_paragraph()
    run_bt = bill_to_heading.add_run('Bill To:')
    run_bt.bold = True
    run_bt.font.size = Pt(12)
    run_bt.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    company_para = doc.add_paragraph()
    company_run = company_para.add_run('<Company>')
    company_run.font.size = Pt(14)
    company_run.bold = True

    # Blank line
    doc.add_paragraph()

    # Invoice details table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    # Row 1: Invoice Number
    table.cell(0, 0).text = 'Invoice Number:'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = '<InvoiceNum>'

    # Row 2: Amount Due
    table.cell(1, 0).text = 'Amount Due:'
    table.cell(1, 0).paragraphs[0].runs[0].bold = True
    table.cell(1, 1).text = '<Amount>'

    # Row 3: Due Date
    table.cell(2, 0).text = 'Due Date:'
    table.cell(2, 0).paragraphs[0].runs[0].bold = True
    table.cell(2, 1).text = '<DueDate>'

    # Blank lines
    doc.add_paragraph()
    doc.add_paragraph()

    # Payment terms
    terms_heading = doc.add_paragraph()
    run_th = terms_heading.add_run('Payment Terms:')
    run_th.bold = True
    run_th.font.size = Pt(11)

    terms = doc.add_paragraph(
        'Payment is due within 30 days of invoice date. '
        'Please remit payment to the bank account details provided separately. '
        'Late payments may incur a 1.5% monthly service charge.'
    )
    terms.paragraph_format.space_after = Pt(6)

    # Footer note
    doc.add_paragraph()
    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fn_run = footer_note.add_run('Thank you for your business!')
    fn_run.italic = True
    fn_run.font.size = Pt(10)
    fn_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT)
    print(f'Template created: {OUTPUT}')


def create_initial():
    create_data_source()
    create_template()

    # Open the template in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
