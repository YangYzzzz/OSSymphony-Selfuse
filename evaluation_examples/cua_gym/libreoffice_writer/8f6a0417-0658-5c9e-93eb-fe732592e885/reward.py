"""
Reward Script: Remove duplicate server hostnames from server_inventory.docx
Task ID: osworld_writer_dedup_005
Domain: libreoffice_writer
Scoring:
  Component 1: No duplicate hostnames remain (0.5 pts)
  Component 2: Correct unique hostname count == 40 (0.2 pts)
  Component 3: First-occurrence order of hostnames is preserved (0.3 pts)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_dedup_005'

# The expected unique hostnames in first-occurrence order from the initial file
EXPECTED_UNIQUE_HOSTNAMES = [
    'web-server-01.prod',
    'db-server-02.prod',
    'web-server-03.prod',
    'app-server-01.prod',
    'cache-server-01.prod',
    'db-server-01.prod',
    'web-server-02.prod',
    'monitor-01.prod',
    'mail-server-01.prod',
    'ftp-server-01.prod',
    'app-server-02.prod',
    'lb-server-01.prod',
    'proxy-server-01.prod',
    'auth-server-01.prod',
    'storage-node-01.prod',
    'backup-server-01.prod',
    'web-server-04.prod',
    'db-server-03.prod',
    'log-server-01.prod',
    'metrics-server-01.prod',
    'api-gateway-01.prod',
    'scheduler-01.prod',
    'queue-server-01.prod',
    'search-server-01.prod',
    'cdn-node-01.prod',
    'dns-server-01.prod',
    'ntp-server-01.prod',
    'vpn-server-01.prod',
    'firewall-01.prod',
    'bastion-host-01.prod',
    'ci-server-01.prod',
    'registry-server-01.prod',
    'artifact-server-01.prod',
    'report-server-01.prod',
    'etl-server-01.prod',
    'ml-worker-01.prod',
    'staging-web-01.prod',
    'staging-db-01.prod',
    'dev-server-01.prod',
    'test-server-01.prod',
]
EXPECTED_COUNT = 40


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires deduplication of server hostnames in the docx file.
    Initial file: 105 lines with 40 unique hostnames (each appearing 2-4 times)
    Golden file: 40 lines with unique hostnames in first-occurrence order
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract all non-empty lines from the document
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    print(f"INFO: Document has {len(doc.paragraphs)} paragraphs, {len(lines)} non-empty lines")

    # Precondition gate: check file has at least some hostnames
    if len(lines) == 0:
        print("CRITICAL: Document has no content")
        print("REWARD: 0.0")
        return 0.0

    # Detect if this is still the initial state (105 lines with duplicates)
    # The initial file has 105 lines; the golden has 40
    if len(lines) == 105:
        print("INFO: File appears to be in initial state (105 lines with duplicates)")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: No duplicate hostnames remain (0.5 points)
    # This FAILS on initial (all 40 hostnames duplicated) and PASSES on golden (all unique)
    try:
        from collections import Counter
        counts = Counter(lines)
        duplicated = {k: v for k, v in counts.items() if v > 1}
        if len(duplicated) == 0:
            print(f"PASS: Component 1 — No duplicate hostnames found ({len(lines)} unique lines) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — {len(duplicated)} hostnames still have duplicates: "
                  f"{list(duplicated.items())[:3]}...")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Correct unique hostname count == 40 (0.2 points)
    # This FAILS on initial (105 non-unique lines) and PASSES on golden (40 unique lines)
    try:
        if len(lines) == EXPECTED_COUNT:
            print(f"PASS: Component 2 — Correct line count: {len(lines)} (expected {EXPECTED_COUNT}) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 2 — Expected {EXPECTED_COUNT} lines, found {len(lines)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: First-occurrence order of hostnames is preserved (0.3 points)
    # The unique hostnames should appear in the same order as first seen in the initial file
    try:
        if lines == EXPECTED_UNIQUE_HOSTNAMES:
            print(f"PASS: Component 3 — Hostnames preserved in first-occurrence order (0.3 pts)")
            total_score += 0.3
        else:
            # Check how many match in order
            match_count = sum(1 for a, b in zip(lines, EXPECTED_UNIQUE_HOSTNAMES) if a == b)
            # Also check if all expected hostnames are present (even if order differs)
            present = set(lines) == set(EXPECTED_UNIQUE_HOSTNAMES)
            if present and len(lines) == EXPECTED_COUNT:
                print(f"FAIL: Component 3 — All {EXPECTED_COUNT} hostnames present but ORDER differs. "
                      f"{match_count}/{EXPECTED_COUNT} in correct position. "
                      f"Expected first: {EXPECTED_UNIQUE_HOSTNAMES[:3]}, "
                      f"Found first: {lines[:3]}")
            else:
                # Find missing/extra hostnames
                expected_set = set(EXPECTED_UNIQUE_HOSTNAMES)
                actual_set = set(lines)
                missing = expected_set - actual_set
                extra = actual_set - expected_set
                print(f"FAIL: Component 3 — Hostname set differs. "
                      f"Missing: {list(missing)[:3]}, Extra: {list(extra)[:3]}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
