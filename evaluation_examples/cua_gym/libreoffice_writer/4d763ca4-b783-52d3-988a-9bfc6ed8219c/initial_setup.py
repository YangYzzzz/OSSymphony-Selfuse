"""
Initial Setup: Investment Prospectus - No sections exist yet
Task ID: writer_struct_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_047'
OUTPUT = f'{WORKDIR}/Desktop/investment_prospectus.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Document title
    title = doc.add_heading('Meridian Capital Partners', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('Investment Prospectus — Series B Offering', level=2)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # First paragraph (part of what should become "Executive Summary")
    p1 = doc.add_paragraph(
        'Meridian Capital Partners presents a compelling opportunity to participate in the Series B '
        'funding round for a high-growth fintech platform operating across Southeast Asia. The fund '
        'targets a minimum return of 3.5x over a seven-year investment horizon, backed by a proven '
        'management team with over $2.4 billion in prior exits. This offering is available exclusively '
        'to accredited investors and qualified institutional buyers as defined under Regulation D of the '
        'Securities Act of 1933.'
    )

    # Second paragraph (part of what should become "Executive Summary")
    p2 = doc.add_paragraph(
        'The investment vehicle is structured as a limited partnership, with Meridian Capital Partners GP LLC '
        'serving as the General Partner. Total target raise for this round is $75 million, with a hard cap '
        'of $100 million. The minimum individual commitment is $250,000. Funds will be deployed across '
        'three primary portfolio companies: PayStream Asia Ltd., ClearLedger Technologies Pte., and '
        'VaultRoute Financial Services, all headquartered in Singapore.'
    )

    # Section divider heading
    doc.add_heading('Company Overview', level=1)

    doc.add_paragraph(
        'PayStream Asia Ltd. was founded in 2019 and provides real-time payment infrastructure to '
        'over 1,200 enterprise clients in six countries. The company processed $18.7 billion in '
        'transaction volume in FY2024, representing a 62% year-over-year growth rate. Net revenue '
        'retention stands at 138%, indicating strong product-market fit and low churn among '
        'enterprise segments.'
    )

    doc.add_paragraph(
        'ClearLedger Technologies Pte. focuses on automated reconciliation and compliance reporting '
        'for regional banks and insurance carriers. Its proprietary ML engine reduces manual '
        'reconciliation effort by up to 87% and integrates with 34 core banking systems. As of Q1 '
        '2025, ClearLedger serves 47 financial institutions with an average contract value of $380,000 '
        'per annum.'
    )

    doc.add_paragraph(
        'VaultRoute Financial Services delivers cross-border remittance and treasury management '
        'capabilities to SMEs and multinational corporations. The platform handles over 800,000 '
        'transactions per month across 12 currency corridors. VaultRoute holds licenses in Singapore, '
        'Malaysia, Indonesia, and the Philippines, and is in the process of securing an Australian '
        'Financial Services License.'
    )

    doc.add_heading('Market Opportunity', level=1)

    doc.add_paragraph(
        'Southeast Asia\'s digital financial services market is projected to reach $1.1 trillion in '
        'gross transaction value by 2030, driven by rapid smartphone adoption, expanding middle-class '
        'populations, and regulatory initiatives supporting open banking frameworks. According to a '
        '2024 report by Bain & Company, fintech penetration in the region remains below 35% across '
        'most product categories, suggesting significant headroom for growth.'
    )

    doc.add_paragraph(
        'The cross-border payments segment alone is expected to grow at a CAGR of 14.3% through '
        '2030, reaching $340 billion in annual flow. Regulatory clarity in Singapore and Malaysia—'
        'including the Monetary Authority of Singapore\'s Payment Services Act revisions—has created '
        'a favorable licensing environment that our portfolio companies have leveraged to establish '
        'first-mover advantages in key corridors.'
    )

    doc.add_heading('Investment Strategy', level=1)

    doc.add_paragraph(
        'Meridian\'s investment thesis centers on backing technology-enabled financial services companies '
        'with defensible regulatory moats, proven unit economics, and multi-country expansion '
        'potential. The General Partner applies a structured due diligence framework, typically '
        'involving 90-day operational audits, third-party technology assessments, and independent '
        'legal reviews across all jurisdictions of operation.'
    )

    doc.add_paragraph(
        'Portfolio construction targets a concentrated strategy: 3 to 5 companies per fund, with '
        'initial positions of $15–$25 million and reserved capital for follow-on rounds. The fund '
        'maintains observer rights on all portfolio company boards and requires quarterly financial '
        'reporting, annual audited statements, and key management retention agreements.'
    )

    doc.add_heading('Risk Factors', level=1)

    doc.add_paragraph(
        'An investment in this fund involves a high degree of risk. Past performance is not indicative '
        'of future results. Key risk factors include, but are not limited to: regulatory changes in '
        'target jurisdictions, currency exchange rate fluctuations, concentration risk given the '
        'limited number of portfolio companies, illiquidity of interests in the fund, and dependence '
        'on key personnel at both the GP and portfolio company levels.'
    )

    doc.add_paragraph(
        'Prospective investors should carefully review the full risk factors section in the Private '
        'Placement Memorandum (PPM) and consult independent legal, tax, and financial advisors before '
        'making any investment decision. This prospectus does not constitute an offer to sell or a '
        'solicitation to buy in any jurisdiction where such offer or solicitation is unlawful.'
    )

    doc.add_heading('Financial Projections', level=1)

    doc.add_paragraph(
        'Based on current portfolio company trajectories and our proprietary benchmarking model, '
        'Meridian projects an aggregate portfolio valuation of $485–$620 million at fund maturity '
        '(Year 7), implying a blended return multiple of 4.85x–6.20x on invested capital under the '
        'base case scenario. The IRR sensitivity analysis indicates a range of 24–38% depending on '
        'exit timing and strategic acquirer availability.'
    )

    doc.add_paragraph(
        'The waterfall structure calls for a preferred return of 8% per annum to LPs on contributed '
        'capital, followed by a catch-up provision of 100% to the GP until a 20% carried interest '
        'split is achieved, thereafter 80/20 LP/GP on all remaining distributions. Management fees '
        'are set at 1.75% per annum during the investment period (Years 1–4) and 1.25% per annum '
        'during the harvest period (Years 5–7).'
    )

    doc.add_heading('Governance & Compliance', level=1)

    doc.add_paragraph(
        'The fund is registered in the Cayman Islands as an exempted limited partnership and is '
        'managed from Singapore. The GP holds a Capital Markets Services license issued by the '
        'Monetary Authority of Singapore (MAS). The fund complies with FATCA and CRS reporting '
        'requirements and maintains robust AML/KYC procedures for all investor onboarding.'
    )

    doc.add_paragraph(
        'An independent advisory board comprising three non-executive directors provides oversight '
        'on valuation, conflict-of-interest matters, and LP communication. Annual audits are '
        'conducted by PricewaterhouseCoopers Singapore, and the fund administrator is Citco Fund '
        'Services (Singapore) Pte. Ltd.'
    )

    doc.add_heading('Contact & Subscription', level=1)

    doc.add_paragraph(
        'For subscription documents, investor questionnaires, and the full Private Placement '
        'Memorandum, please contact our investor relations team at ir@meridiancapitalpartners.com '
        'or call +65 6389 7200 during Singapore business hours (9:00 AM–6:00 PM SGT, Monday through '
        'Friday). Hard copies may be requested by writing to our registered office at 1 Raffles Place, '
        '#42-01 One Raffles Place Tower 2, Singapore 048616.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
