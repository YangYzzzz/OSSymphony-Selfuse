"""
Initial Setup: Remove the existing table of contents from a document
Task ID: writer_mt_061
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_061'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_field(doc):
    """Add a Table of Contents field to the document using field codes."""
    # Add a "Table of Contents" title
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = toc_title.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"

    # Add an empty line after title
    doc.add_paragraph()

    # Create TOC field using XML field codes
    paragraph = doc.add_paragraph()
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar2)

    # Add static TOC entries that simulate what the TOC would render
    toc_entries = []

    # We'll collect entry info first, then add them
    sections_data = get_sections_data()
    entry_num = 0
    for section in sections_data:
        entry_num += 1
        level = section['level']
        title = section['title']
        indent = Inches(0.25 * (level - 1))
        toc_entries.append((title, level, indent, entry_num))

    # Add TOC entry paragraphs between separate and end field chars
    for title, level, indent, page_num in toc_entries:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.left_indent = indent
        toc_para.paragraph_format.space_before = Pt(2)
        toc_para.paragraph_format.space_after = Pt(2)
        toc_run = toc_para.add_run(f"{title}")
        toc_run.font.size = Pt(11) if level == 1 else Pt(10)
        if level == 1:
            toc_run.bold = True
        # Add tab and page number
        tab_run = toc_para.add_run(f"\t{page_num}")
        tab_run.font.size = Pt(10)

    # Add another empty line
    doc.add_paragraph()

    # End the TOC field
    end_para = doc.add_paragraph()
    run_end = end_para.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fldChar3)

    # Page break after TOC
    doc.add_page_break()


def get_sections_data():
    """Return the content structure for the proposal document."""
    sections = [
        {'level': 1, 'title': 'Executive Summary', 'paragraphs': [
            "Meridian Analytics proposes the development and deployment of a comprehensive data analytics platform designed to transform how regional healthcare providers manage patient outcomes and operational efficiency. This initiative addresses the growing demand for data-driven decision making in healthcare administration.",
            "Our proposed solution leverages machine learning algorithms and real-time data processing to deliver actionable insights across clinical operations, financial management, and patient engagement metrics. The platform will integrate seamlessly with existing Electronic Health Record (EHR) systems and comply with all HIPAA regulations.",
            "The total investment required is $2.4 million over a 24-month implementation timeline, with projected annual savings of $3.8 million by Year 3 of full deployment. This proposal outlines the technical architecture, implementation phases, risk mitigation strategies, and expected return on investment."
        ]},
        {'level': 1, 'title': 'Background and Problem Statement', 'paragraphs': [
            "Healthcare organizations across the Pacific Northwest are facing unprecedented challenges in data management. With patient volumes increasing by 12% annually and regulatory requirements becoming more complex, traditional data management approaches are no longer sufficient.",
            "Current systems operate in silos, with clinical data, financial records, and operational metrics stored in disconnected databases. This fragmentation leads to delayed reporting, inconsistent data quality, and missed opportunities for proactive patient care interventions.",
            "A 2025 survey conducted by the Healthcare Information and Management Systems Society (HIMSS) found that 67% of mid-sized healthcare providers cited data integration as their primary operational challenge. Furthermore, organizations that adopted unified analytics platforms reported 23% improvements in patient readmission rates."
        ]},
        {'level': 2, 'title': 'Current Infrastructure Assessment', 'paragraphs': [
            "Our preliminary assessment of the client's existing infrastructure reveals a heterogeneous technology landscape comprising three primary EHR systems (Epic, Cerner, and Meditech), two legacy billing platforms, and numerous departmental databases built on Microsoft Access and FileMaker Pro.",
            "The client currently employs 14 full-time data analysts across five departments, each using different tools and methodologies. Report generation typically requires 3-5 business days, and cross-departmental analysis is performed manually using spreadsheets exported from individual systems."
        ]},
        {'level': 2, 'title': 'Key Pain Points Identified', 'paragraphs': [
            "Through stakeholder interviews and workflow analysis, we identified the following critical pain points: (1) average report turnaround time of 4.2 days, (2) data inconsistency rates of approximately 8.3% across systems, (3) inability to perform real-time monitoring of key performance indicators, and (4) compliance reporting requires manual compilation from six different source systems.",
            "These inefficiencies result in an estimated annual cost of $1.9 million in staff overtime, redundant data entry, and delayed decision-making. The risk of regulatory non-compliance adds an additional projected liability of $500,000 annually."
        ]},
        {'level': 1, 'title': 'Proposed Solution', 'paragraphs': [
            "Meridian Analytics proposes the implementation of the HealthPulse Analytics Platform, a cloud-native solution built on a modern data lakehouse architecture. The platform will serve as the single source of truth for all clinical, financial, and operational data across the organization.",
            "HealthPulse leverages Apache Spark for distributed data processing, Delta Lake for ACID-compliant data storage, and a custom-built visualization layer powered by D3.js and React. The machine learning pipeline utilizes TensorFlow and scikit-learn for predictive analytics, with models trained on de-identified historical patient data."
        ]},
        {'level': 2, 'title': 'Technical Architecture', 'paragraphs': [
            "The platform architecture consists of four primary layers: (1) Data Ingestion Layer utilizing Apache Kafka for real-time streaming and Apache NiFi for batch processing, (2) Storage Layer built on Delta Lake with automatic schema evolution, (3) Processing Layer powered by Apache Spark with dedicated clusters for ETL, analytics, and ML workloads, and (4) Presentation Layer delivering interactive dashboards and automated reports.",
            "All data in transit will be encrypted using TLS 1.3, and data at rest will employ AES-256 encryption. The platform will be deployed on AWS GovCloud to meet federal compliance requirements, with multi-region failover capabilities ensuring 99.97% uptime."
        ]},
        {'level': 2, 'title': 'Integration Strategy', 'paragraphs': [
            "Integration with existing EHR systems will be accomplished through HL7 FHIR APIs, which provide standardized endpoints for clinical data exchange. For legacy systems that do not support FHIR, we will deploy custom ETL connectors using Apache Camel integration framework.",
            "The integration timeline allocates four weeks per major system connection, with parallel workstreams for non-dependent systems. A dedicated integration testing environment will mirror the production topology to validate data flows before go-live."
        ]},
        {'level': 2, 'title': 'Machine Learning Models', 'paragraphs': [
            "The platform will include five pre-trained ML models: (1) Patient Readmission Risk Predictor using gradient-boosted decision trees, (2) Revenue Cycle Anomaly Detector using isolation forests, (3) Staff Scheduling Optimizer using constraint satisfaction, (4) Clinical Pathway Recommender using collaborative filtering, and (5) Supply Chain Demand Forecaster using LSTM neural networks.",
            "Each model will undergo quarterly retraining using the latest 18 months of data, with automated drift detection triggering ad-hoc retraining when model performance degrades below established thresholds. All model predictions will include confidence intervals and explainability scores generated by SHAP values."
        ]},
        {'level': 1, 'title': 'Implementation Plan', 'paragraphs': [
            "The implementation will follow an agile methodology organized into six quarterly phases over the 24-month project timeline. Each phase will deliver functional increments that provide immediate value while building toward the complete platform vision."
        ]},
        {'level': 2, 'title': 'Phase 1: Foundation (Months 1-4)', 'paragraphs': [
            "The foundation phase focuses on infrastructure provisioning, security configuration, and initial data pipeline development. Key deliverables include AWS GovCloud environment setup, network architecture implementation, CI/CD pipeline configuration, and integration with the primary EHR system (Epic).",
            "During this phase, we will also conduct comprehensive data profiling across all source systems to establish baseline data quality metrics and identify transformation requirements. A data governance framework will be developed in collaboration with the client's compliance and IT security teams."
        ]},
        {'level': 2, 'title': 'Phase 2: Core Analytics (Months 5-8)', 'paragraphs': [
            "The core analytics phase delivers the first production dashboards and automated reports. Priority use cases include daily census reporting, revenue cycle monitoring, emergency department throughput analysis, and surgical scheduling optimization.",
            "This phase also includes the deployment of the first two ML models: Patient Readmission Risk Predictor and Revenue Cycle Anomaly Detector. Both models will enter a 30-day shadow mode, running in parallel with existing processes to validate accuracy before full production deployment."
        ]},
        {'level': 2, 'title': 'Phase 3: Advanced Integration (Months 9-12)', 'paragraphs': [
            "Phase 3 extends data integration to remaining EHR systems (Cerner and Meditech) and legacy billing platforms. Cross-system data reconciliation processes will be automated, and historical data migration will begin for the unified analytics warehouse.",
            "Additional dashboards will cover quality metrics required for CMS reporting, infection control surveillance, and medication management analytics. The Staff Scheduling Optimizer ML model will be deployed during this phase."
        ]},
        {'level': 2, 'title': 'Phase 4: Predictive Analytics (Months 13-16)', 'paragraphs': [
            "The predictive analytics phase introduces the Clinical Pathway Recommender and Supply Chain Demand Forecaster models. These models require the integrated dataset from Phase 3 to achieve acceptable accuracy levels.",
            "This phase also delivers a self-service analytics portal allowing department managers to create custom reports and visualizations without requiring data analyst support. The portal includes a natural language query interface powered by a fine-tuned GPT model."
        ]},
        {'level': 2, 'title': 'Phase 5: Optimization (Months 17-20)', 'paragraphs': [
            "Phase 5 focuses on performance optimization, advanced automation, and user adoption acceleration. Platform response times will be tuned to sub-second latency for interactive queries and under 30 seconds for complex analytical workloads.",
            "Automated alerting and notification systems will be deployed, enabling proactive identification of clinical and operational anomalies. Integration with mobile platforms will allow executives and clinical leaders to access key metrics from smartphones and tablets."
        ]},
        {'level': 2, 'title': 'Phase 6: Maturation (Months 21-24)', 'paragraphs': [
            "The final phase focuses on knowledge transfer, documentation, and establishing ongoing operational procedures. The client's internal analytics team will receive comprehensive training covering platform administration, model retraining, and dashboard development.",
            "A 90-day hypercare period will follow the formal handoff, during which Meridian Analytics will provide dedicated support with guaranteed 2-hour response times for critical issues and 8-hour response for non-critical inquiries."
        ]},
        {'level': 1, 'title': 'Risk Assessment and Mitigation', 'paragraphs': [
            "Our risk management framework identifies and addresses potential challenges across technical, organizational, and regulatory dimensions. Each risk has been evaluated using a probability-impact matrix, and mitigation strategies have been developed for all high and medium-rated risks."
        ]},
        {'level': 2, 'title': 'Technical Risks', 'paragraphs': [
            "Data quality inconsistencies across legacy systems represent the highest technical risk (probability: high, impact: high). Mitigation includes automated data quality scoring, quarantine workflows for suspect records, and a dedicated data stewardship program with subject matter experts from each clinical department.",
            "System performance under peak loads presents a medium risk. Our architecture employs auto-scaling compute clusters and query optimization techniques including materialized views, partition pruning, and adaptive query execution. Load testing will simulate 3x expected peak concurrent users."
        ]},
        {'level': 2, 'title': 'Organizational Risks', 'paragraphs': [
            "Change resistance from staff accustomed to existing workflows is anticipated. Our change management approach includes executive sponsorship engagement, department-level champions, hands-on training workshops, and a phased rollout strategy that allows gradual adoption rather than abrupt transitions.",
            "Staff turnover during the implementation period could impact institutional knowledge. We mitigate this through comprehensive documentation, cross-training programs, and recorded training sessions that serve as an ongoing reference library."
        ]},
        {'level': 2, 'title': 'Regulatory Compliance Risks', 'paragraphs': [
            "HIPAA compliance is paramount and non-negotiable. Our platform has been designed from the ground up with privacy-by-design principles: role-based access control, audit logging of all data access, automatic de-identification for research datasets, and encryption at every layer.",
            "We will engage a third-party HIPAA auditor to conduct a pre-deployment compliance assessment and annual reviews thereafter. Business Associate Agreements will be established with all cloud infrastructure providers and third-party services involved in data processing."
        ]},
        {'level': 1, 'title': 'Budget and Financial Projections', 'paragraphs': [
            "The total project investment of $2.4 million is distributed across personnel costs (58%), infrastructure and licensing (27%), and contingency reserves (15%). Detailed line-item budgets are provided in Appendix A."
        ]},
        {'level': 2, 'title': 'Cost Breakdown', 'paragraphs': [
            "Personnel costs include a dedicated team of 8 specialists: Project Manager ($185,000), Lead Data Architect ($195,000), 2 Data Engineers ($170,000 each), ML Engineer ($180,000), Frontend Developer ($160,000), QA/Test Engineer ($145,000), and DevOps Engineer ($155,000). These represent fully-loaded annual rates including benefits and overhead.",
            "Infrastructure costs encompass AWS GovCloud compute and storage ($340,000 annually), software licensing for monitoring and security tools ($145,000 annually), and development environment provisioning ($65,000 one-time). The contingency reserve of $360,000 provides a buffer for scope adjustments and unforeseen technical challenges."
        ]},
        {'level': 2, 'title': 'Return on Investment', 'paragraphs': [
            "Based on conservative estimates, the platform will generate the following annual savings by Year 3: reduced staff overtime in manual reporting ($680,000), eliminated redundant data entry across departments ($420,000), decreased patient readmission penalties ($890,000), improved revenue cycle efficiency ($1,250,000), and reduced compliance audit preparation costs ($560,000).",
            "The total projected annual savings of $3.8 million represents a 158% return on the initial investment. The break-even point is projected at 18 months post-full-deployment, or 42 months from project inception."
        ]},
        {'level': 1, 'title': 'Team Qualifications', 'paragraphs': [
            "Meridian Analytics has successfully delivered 23 healthcare analytics implementations over the past seven years, serving organizations ranging from 200-bed community hospitals to 2,000-bed academic medical centers. Our team members hold relevant certifications including AWS Solutions Architect Professional, Google Cloud Professional Data Engineer, CHDA (Certified Health Data Analyst), and PMP."
        ]},
        {'level': 2, 'title': 'Key Personnel', 'paragraphs': [
            "Dr. Elena Vasquez, Project Director, brings 15 years of experience in healthcare informatics with a PhD in Biomedical Informatics from Stanford University. She has led the implementation of analytics platforms at Kaiser Permanente, Cleveland Clinic, and Mayo Clinic.",
            "James Okonkwo, Lead Data Architect, holds dual master's degrees in Computer Science and Public Health from MIT. His expertise in FHIR interoperability standards and healthcare data modeling has been recognized through published research in the Journal of Biomedical Informatics and invited presentations at HIMSS and AMIA conferences.",
            "Sarah Kim, ML Engineering Lead, previously served as Senior Data Scientist at Google Health, where she developed predictive models for clinical decision support. Her patient readmission prediction model achieved an AUC of 0.89 in production across 12 hospital systems."
        ]},
        {'level': 1, 'title': 'Conclusion and Next Steps', 'paragraphs': [
            "The HealthPulse Analytics Platform represents a transformative investment in data-driven healthcare management. By consolidating fragmented data sources, automating routine analyses, and deploying predictive capabilities, the platform will position the organization as a regional leader in healthcare analytics.",
            "We recommend proceeding with a two-week discovery engagement to finalize technical requirements, establish governance frameworks, and develop a detailed project schedule. This discovery phase requires a commitment of $45,000 and will produce a comprehensive Statement of Work for the full implementation.",
            "Meridian Analytics is confident in our ability to deliver exceptional results and looks forward to the opportunity to partner with your organization on this important initiative. Please direct any questions regarding this proposal to Dr. Elena Vasquez at evasquez@meridiananalytics.com or (503) 555-0142."
        ]},
    ]
    return sections


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Document Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run("Draft Proposal")
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("HealthPulse Analytics Platform")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.font.name = "Calibri"

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_after = Pt(24)
    run = author_para.add_run("Prepared by Meridian Analytics Inc.\nMarch 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = "Calibri"

    # --- Table of Contents ---
    add_toc_field(doc)

    # --- Document Body Content ---
    sections = get_sections_data()
    for section in sections:
        level = section['level']
        title = section['title']

        # Add heading
        heading = doc.add_heading(title, level=level)

        # Add body paragraphs
        for para_text in section['paragraphs']:
            para = doc.add_paragraph(para_text)
            para.paragraph_format.space_after = Pt(8)

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
