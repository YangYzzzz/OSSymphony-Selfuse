"""
Reward Script: Remove the existing table of contents from the document
Task ID: writer_mt_061
Domain: libreoffice_writer
Scoring:
  Component 1 — TOC field codes removed (0.3 pts)
  Component 2 — TOC text content removed, first heading near doc start (0.4 pts)
  Component 3 — Body content preserved: all headings intact (0.3 pts)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_061'

# --- Persistence hook (save any unsaved LibreOffice changes) ---
def persist_app_state():
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")

# --- Expected headings (from initial doc body content) ---
EXPECTED_HEADINGS = [
    "Executive Summary",
    "Background and Problem Statement",
    "Current Infrastructure Assessment",
    "Key Pain Points Identified",
    "Proposed Solution",
    "Technical Architecture",
    "Integration Strategy",
    "Machine Learning Models",
    "Implementation Plan",
    "Phase 1: Foundation (Months 1-4)",
    "Phase 2: Core Analytics (Months 5-8)",
    "Phase 3: Advanced Integration (Months 9-12)",
    "Phase 4: Predictive Analytics (Months 13-16)",
    "Phase 5: Optimization (Months 17-20)",
    "Phase 6: Maturation (Months 21-24)",
    "Risk Assessment and Mitigation",
    "Technical Risks",
    "Organizational Risks",
    "Regulatory Compliance Risks",
    "Budget and Financial Projections",
    "Cost Breakdown",
    "Return on Investment",
    "Team Qualifications",
    "Key Personnel",
    "Conclusion and Next Steps",
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0
    toc_field_found = True   # assume TOC present until proven otherwise
    toc_title_found = True   # assume TOC title present until proven otherwise

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: TOC field codes removed (0.3 points)
    # In the initial doc, there is a TOC field instruction.
    # In the golden doc, it should be completely absent.
    try:
        toc_field_found = False
        for para in doc.paragraphs:
            for run in para.runs:
                for fld in run.element.findall('.//w:instrText', ns):
                    if fld.text and 'TOC' in fld.text:
                        toc_field_found = True
                        break
            if toc_field_found:
                break

        # Also check body-level XML for TOC field codes (may be outside runs)
        if not toc_field_found:
            body_xml = doc.element.body.xml if hasattr(doc.element.body, 'xml') else ''
            if 'TOC' in body_xml and 'instrText' in body_xml:
                # Double check: parse for actual instrText containing TOC
                for instr in doc.element.body.findall('.//w:instrText', ns):
                    if instr.text and 'TOC' in instr.text:
                        toc_field_found = True
                        break

        if not toc_field_found:
            print("PASS: Component 1 — No TOC field codes found (0.3 pts)")
            total_score += 0.3
        else:
            print("FAIL: Component 1 — TOC field code still present in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: TOC text content removed (0.4 points)
    # In the initial doc, "Table of Contents" appears as a heading/text before
    # the actual headings. After removal, the first heading-styled paragraph
    # should appear early in the document (within first ~10 paragraphs).
    # Also no paragraph should contain "Table of Contents" as standalone text.
    try:
        # Check no "Table of Contents" standalone paragraph
        toc_title_found = False
        for p in doc.paragraphs[:20]:
            if p.text.strip().lower() == "table of contents":
                toc_title_found = True
                break

        # Check that the first Heading paragraph appears within the first 10 paragraphs
        first_heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.style and 'Heading' in p.style.name:
                first_heading_idx = i
                break

        toc_entries_found = False
        # Check for TOC-like entries (text followed by tab and page number)
        # These appear as "Some Heading\t5" in the initial doc's TOC section
        import re
        for p in doc.paragraphs[:20]:
            text = p.text.strip()
            # TOC entries: heading text followed by tab and a number
            if re.match(r'^.+\t\d+$', text) and (p.style is None or 'Heading' not in p.style.name):
                toc_entries_found = True
                break

        comp2_pass = (
            not toc_title_found
            and first_heading_idx is not None
            and first_heading_idx < 10
            and not toc_entries_found
        )

        if comp2_pass:
            print(f"PASS: Component 2 — TOC text removed. First heading at paragraph [{first_heading_idx}] (0.4 pts)")
            total_score += 0.4
        else:
            reasons = []
            if toc_title_found:
                reasons.append("'Table of Contents' paragraph still present")
            if first_heading_idx is None:
                reasons.append("No heading paragraphs found at all")
            elif first_heading_idx >= 10:
                reasons.append(f"First heading at paragraph [{first_heading_idx}], expected < 10")
            if toc_entries_found:
                reasons.append("TOC-like entries (text + tab + page number) still present")
            print(f"FAIL: Component 2 — {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: TOC removed AND body content preserved (0.3 points)
    # Compound check: TOC must be removed (task change) AND all headings intact (no collateral damage).
    # This only awards points when the TOC is already gone, so it FAILS on initial_env.
    try:
        doc_headings = []
        for p in doc.paragraphs:
            if p.style and 'Heading' in p.style.name:
                doc_headings.append(p.text.strip())

        # Check that all expected headings are present
        missing_headings = []
        for expected in EXPECTED_HEADINGS:
            if expected not in doc_headings:
                missing_headings.append(expected)

        headings_ok = len(missing_headings) == 0 and len(doc_headings) == len(EXPECTED_HEADINGS)

        # Anchor to task change: TOC must also be removed (no TOC field codes + no TOC title)
        toc_removed = not toc_field_found and not toc_title_found

        if toc_removed and headings_ok:
            print(f"PASS: Component 3 — TOC removed AND all {len(EXPECTED_HEADINGS)} headings intact (0.3 pts)")
            total_score += 0.3
        else:
            reasons = []
            if not toc_removed:
                reasons.append("TOC still present (precondition for this component)")
            if missing_headings:
                reasons.append(f"Missing headings: {missing_headings[:5]}...")
            if len(doc_headings) != len(EXPECTED_HEADINGS):
                reasons.append(f"Expected {len(EXPECTED_HEADINGS)} headings, found {len(doc_headings)}")
            print(f"FAIL: Component 3 — {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
