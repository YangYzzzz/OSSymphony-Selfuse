"""
Initial Setup: Legal document with standard autocorrect settings
Task ID: writer_legal_088
Domain: libreoffice_writer

Creates a legal document and ensures no custom autocorrect entries exist
for Plf, Def, Jdg, Mtn, Stip.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_088'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Document title
    title = doc.add_heading('MEMORANDUM OF LAW IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Case caption
    caption = doc.add_paragraph()
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption.add_run('IN THE SUPERIOR COURT OF COLUMBIA COUNTY')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    caption2 = doc.add_paragraph()
    caption2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = caption2.add_run('CIVIL DIVISION')
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Case number
    case_no = doc.add_paragraph()
    case_no.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_cn = case_no.add_run('Case No. 2025-CV-04821')
    run_cn.font.size = Pt(11)
    run_cn.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    # Parties
    parties = doc.add_paragraph()
    parties.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r1 = parties.add_run('HENDERSON PROPERTIES, LLC')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    parties.add_run('\n')
    r2 = parties.add_run('Plaintiff,')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)
    r2.italic = True
    parties.add_run('\n')
    r3 = parties.add_run('v.')
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(11)
    parties.add_run('\n')
    r4 = parties.add_run('WESTBROOK CONSTRUCTION, INC.')
    r4.bold = True
    r4.font.name = 'Times New Roman'
    r4.font.size = Pt(11)
    parties.add_run('\n')
    r5 = parties.add_run('Defendant.')
    r5.font.name = 'Times New Roman'
    r5.font.size = Pt(11)
    r5.italic = True

    doc.add_paragraph()  # spacer

    # Section I
    h2 = doc.add_heading('I. INTRODUCTION', level=2)

    intro = doc.add_paragraph()
    intro.paragraph_format.line_spacing = 1.5
    intro.paragraph_format.first_line_indent = Inches(0.5)
    run_intro = intro.add_run(
        'Plaintiff Henderson Properties, LLC ("Henderson") respectfully submits this '
        'Memorandum of Law in Support of its Motion for Summary Judgment against Defendant '
        'Westbrook Construction, Inc. ("Westbrook"). The undisputed material facts establish '
        'that Westbrook breached the construction agreement dated March 14, 2024, by failing '
        'to complete the renovation of the commercial property located at 1847 Oakridge Boulevard, '
        'Columbia, within the contractually specified timeframe.'
    )
    run_intro.font.name = 'Times New Roman'
    run_intro.font.size = Pt(11)

    intro2 = doc.add_paragraph()
    intro2.paragraph_format.line_spacing = 1.5
    intro2.paragraph_format.first_line_indent = Inches(0.5)
    run_intro2 = intro2.add_run(
        'As demonstrated below, the Plaintiff is entitled to judgment as a matter of law on its '
        'claims for breach of contract and recovery of liquidated damages in the amount of $287,500.00.'
    )
    run_intro2.font.name = 'Times New Roman'
    run_intro2.font.size = Pt(11)

    # Section II
    doc.add_heading('II. STATEMENT OF UNDISPUTED FACTS', level=2)

    facts = [
        'On March 14, 2024, Henderson and Westbrook entered into a Construction Services Agreement '
        '("the Agreement") for the renovation of the commercial property at 1847 Oakridge Boulevard. '
        '(Ex. A, Agreement at 1.)',

        'The Agreement specified a completion date of September 30, 2024, with liquidated damages '
        'of $2,500.00 per day for each day of delay beyond the completion date. (Ex. A, Agreement '
        'at 7, Section 12.3.)',

        'Westbrook commenced work on April 8, 2024, and was provided full access to the property '
        'as required under Section 4.1 of the Agreement. (Ex. B, Site Access Log.)',

        'As of January 15, 2025, Westbrook had completed only approximately 62% of the contracted '
        'work scope, as documented in the independent inspection report prepared by Marshall Engineering '
        'Associates. (Ex. C, Inspection Report at 3-4.)',

        'Henderson issued three formal notices of default to Westbrook on October 15, 2024, November 12, '
        '2024, and December 3, 2024, each sent via certified mail to Westbrook\'s registered address. '
        '(Ex. D, Default Notices and Certified Mail Receipts.)',

        'Westbrook failed to cure the defaults within the thirty-day cure period specified in Section '
        '15.2 of the Agreement. (Ex. E, Correspondence File.)',

        'Henderson terminated the Agreement effective January 20, 2025, pursuant to Section 15.4, '
        'and engaged Morrison & Blake General Contractors to complete the remaining work at a cost '
        'of $412,000.00. (Ex. F, Termination Letter; Ex. G, Morrison & Blake Contract.)',
    ]

    for i, fact in enumerate(facts, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Inches(0.5)
        run_num = p.add_run(f'{i}. ')
        run_num.bold = True
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(11)
        run_fact = p.add_run(fact)
        run_fact.font.name = 'Times New Roman'
        run_fact.font.size = Pt(11)

    # Section III
    doc.add_heading('III. ARGUMENT', level=2)

    arg_a = doc.add_heading('A. Standard of Review', level=3)

    std_review = doc.add_paragraph()
    std_review.paragraph_format.line_spacing = 1.5
    std_review.paragraph_format.first_line_indent = Inches(0.5)
    run_sr = std_review.add_run(
        'Summary judgment is appropriate when "there is no genuine dispute as to any material '
        'fact and the movant is entitled to judgment as a matter of law." Columbia R. Civ. P. 56(a). '
        'The moving party bears the initial burden of demonstrating the absence of a genuine issue '
        'of material fact. Once met, the burden shifts to the nonmoving party to set forth specific '
        'facts showing a genuine issue for trial. Anderson v. Liberty Lobby, Inc., 477 U.S. 242, 248 (1986).'
    )
    run_sr.font.name = 'Times New Roman'
    run_sr.font.size = Pt(11)

    arg_b = doc.add_heading('B. Westbrook Breached the Agreement', level=3)

    breach = doc.add_paragraph()
    breach.paragraph_format.line_spacing = 1.5
    breach.paragraph_format.first_line_indent = Inches(0.5)
    run_breach = breach.add_run(
        'To establish breach of contract under Columbia law, the Plaintiff must demonstrate: '
        '(1) the existence of a valid and enforceable contract; (2) performance or tendered '
        'performance by the Plaintiff; (3) breach by the Defendant; and (4) resulting damages. '
        'Reeves v. Sanderson Plumbing Products, Inc., 530 U.S. 133, 150 (2000).'
    )
    run_breach.font.name = 'Times New Roman'
    run_breach.font.size = Pt(11)

    breach2 = doc.add_paragraph()
    breach2.paragraph_format.line_spacing = 1.5
    breach2.paragraph_format.first_line_indent = Inches(0.5)
    run_breach2 = breach2.add_run(
        'Each element is satisfied here. The Agreement is a valid, executed contract. Henderson '
        'performed all obligations, including timely payment of progress invoices totaling $635,000.00. '
        'Westbrook failed to complete the work by September 30, 2024, and as of termination had '
        'completed only 62% of the scope. Henderson has suffered damages including liquidated damages '
        'of $287,500.00 (115 days at $2,500.00/day) and completion costs of $412,000.00.'
    )
    run_breach2.font.name = 'Times New Roman'
    run_breach2.font.size = Pt(11)

    # Section IV
    doc.add_heading('IV. CONCLUSION', level=2)

    conclusion = doc.add_paragraph()
    conclusion.paragraph_format.line_spacing = 1.5
    conclusion.paragraph_format.first_line_indent = Inches(0.5)
    run_conc = conclusion.add_run(
        'For the foregoing reasons, Plaintiff Henderson Properties, LLC respectfully requests '
        'that this Court grant its Motion for Summary Judgment against Defendant Westbrook '
        'Construction, Inc. and enter judgment in the amount of $699,500.00, plus pre-judgment '
        'interest, attorneys\' fees, and costs as permitted by law and the Agreement.'
    )
    run_conc.font.name = 'Times New Roman'
    run_conc.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Signature block
    sig = doc.add_paragraph()
    sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sig.paragraph_format.left_indent = Inches(3.5)
    run_sig = sig.add_run('Respectfully submitted,')
    run_sig.font.name = 'Times New Roman'
    run_sig.font.size = Pt(11)

    sig2 = doc.add_paragraph()
    sig2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sig2.paragraph_format.left_indent = Inches(3.5)
    sig2.paragraph_format.space_before = Pt(36)
    run_sig2 = sig2.add_run('___________________________')
    run_sig2.font.name = 'Times New Roman'
    run_sig2.font.size = Pt(11)

    sig3 = doc.add_paragraph()
    sig3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sig3.paragraph_format.left_indent = Inches(3.5)
    run_sig3 = sig3.add_run('Rachel M. Thornton, Esq.\nBar No. 47823\nThornton & Associates, P.C.\n'
                             '2100 Commerce Street, Suite 800\nColumbia, ST 30301\n'
                             'Tel: (555) 482-1900\nrtthornton@thorntonlaw.com\n'
                             'Counsel for Plaintiff Henderson Properties, LLC')
    run_sig3.font.name = 'Times New Roman'
    run_sig3.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure no custom autocorrect entries exist for our abbreviations
    # The default system acor file doesn't have these entries, so we just
    # need to make sure no user-level override exists
    autocorr_dir = '/home/user/.config/libreoffice/4/user/autocorr'
    user_acor = os.path.join(autocorr_dir, 'acor_en-US.dat')
    if os.path.exists(user_acor):
        # Check if it has our entries and remove them
        import zipfile
        import xml.etree.ElementTree as ET
        import shutil
        import tempfile

        with zipfile.ZipFile(user_acor, 'r') as z:
            doc_list = z.read('DocumentList.xml').decode('utf-8')

        # Parse and check for our entries
        abbrevs = {'Plf', 'Def', 'Jdg', 'Mtn', 'Stip'}
        root = ET.fromstring(doc_list)
        ns = {'bl': 'http://openoffice.org/2001/block-list'}
        entries_to_remove = []
        for block in root.findall('bl:block', ns):
            abbr = block.get('{http://openoffice.org/2001/block-list}abbreviated-name', '')
            if abbr in abbrevs:
                entries_to_remove.append(block)

        if entries_to_remove:
            for block in entries_to_remove:
                root.remove(block)
            new_xml = ET.tostring(root, encoding='unicode', xml_declaration=True)
            # Rebuild the zip
            tmp = user_acor + '.tmp'
            with zipfile.ZipFile(user_acor, 'r') as zin, zipfile.ZipFile(tmp, 'w') as zout:
                for item in zin.infolist():
                    if item.filename == 'DocumentList.xml':
                        zout.writestr(item, new_xml)
                    else:
                        zout.writestr(item, zin.read(item.filename))
            shutil.move(tmp, user_acor)
            print('Removed existing legal abbreviation entries from autocorrect')
        else:
            print('No legal abbreviation entries found in user autocorrect - OK')
    else:
        print(f'No user autocorrect file at {user_acor} - standard settings intact')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
