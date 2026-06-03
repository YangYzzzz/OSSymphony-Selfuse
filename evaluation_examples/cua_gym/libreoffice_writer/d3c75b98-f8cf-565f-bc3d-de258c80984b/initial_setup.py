"""
Initial Setup: Create attendance document with a 5x12 table (auto row heights)
Task ID: writer_tm_015
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Monthly Attendance Record', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle paragraph
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.add_run('Engineering Department - March 2025')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Create 5-column x 12-row table
    # Row 0 = header, Rows 1-11 = data
    table = doc.add_table(rows=12, cols=5)
    table.style = 'Table Grid'

    # Headers
    headers = ['Employee Name', 'Date', 'Time In', 'Time Out', 'Status']
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Shade header row
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '2F5496',
        })
        shading.append(shd)

    # Data rows (11 data rows)
    data = [
        ['Sarah Chen',       '2025-03-03', '08:45', '17:30', 'Present'],
        ['Marcus Johnson',   '2025-03-03', '09:02', '17:15', 'Present'],
        ['Priya Patel',      '2025-03-03', '08:30', '17:00', 'Present'],
        ['James O\'Brien',   '2025-03-03', '',      '',      'Absent'],
        ['Yuki Tanaka',      '2025-03-03', '08:55', '17:45', 'Present'],
        ['Elena Rodriguez',  '2025-03-03', '09:15', '17:30', 'Late'],
        ['David Kim',        '2025-03-03', '08:40', '17:00', 'Present'],
        ['Amara Okafor',     '2025-03-03', '08:50', '17:20', 'Present'],
        ['Thomas Weber',     '2025-03-03', '10:00', '17:30', 'Late'],
        ['Lin Zhang',        '2025-03-03', '08:35', '17:10', 'Present'],
        ['Rachel Foster',    '2025-03-03', '08:48', '17:25', 'Present'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    # Set column widths (approximate)
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(1.2)
        row.cells[2].width = Inches(1.0)
        row.cells[3].width = Inches(1.0)
        row.cells[4].width = Inches(1.0)

    # NOTE: Row heights are left at auto (default) - no explicit height set
    # The task requires setting them to 1.2cm

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
