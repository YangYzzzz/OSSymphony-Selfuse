"""
Reward Script: Set all row heights in attendance table to 1.2cm
Task ID: writer_tm_015
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): All rows have fixed height rule (hRule="exact")
  Component 2 (0.5): All row heights are ~680 twips (1.2cm)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_015'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that all rows in the attendance table have a fixed height of 1.2cm.
    1.2 cm = 1.2 * 567 twips = 680.4 twips. We accept 670-690 as tolerance.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    print(f"INFO: Table has {num_rows} rows x {len(table.columns)} cols")

    # Component 1: All rows have fixed height rule (hRule="exact") — 0.5 points
    try:
        exact_count = 0
        for ri, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is not None:
                h_elem = trPr.find(qn('w:trHeight'))
                if h_elem is not None:
                    hRule = h_elem.get(qn('w:hRule'))
                    if hRule == 'exact':
                        exact_count += 1
                    else:
                        print(f"  Row {ri}: hRule={hRule} (expected 'exact')")
                else:
                    print(f"  Row {ri}: no trHeight element")
            else:
                print(f"  Row {ri}: no trPr element")

        if exact_count == num_rows:
            print(f"PASS: Component 1 — All {num_rows} rows have hRule='exact' (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — {exact_count}/{num_rows} rows have hRule='exact'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All row heights are ~680 twips (1.2cm) — 0.5 points
    # 1.2cm = 680.4 twips. Accept range 670-690 to allow minor rounding.
    try:
        correct_height_count = 0
        TARGET_TWIPS = 680
        TOLERANCE = 10  # accept 670-690

        for ri, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is not None:
                h_elem = trPr.find(qn('w:trHeight'))
                if h_elem is not None:
                    val_str = h_elem.get(qn('w:val'))
                    if val_str is not None:
                        val = int(val_str)
                        if abs(val - TARGET_TWIPS) <= TOLERANCE:
                            correct_height_count += 1
                        else:
                            print(f"  Row {ri}: height={val} twips (expected ~{TARGET_TWIPS})")
                    else:
                        print(f"  Row {ri}: trHeight has no val attribute")
                else:
                    print(f"  Row {ri}: no trHeight element")
            else:
                print(f"  Row {ri}: no trPr element")

        if correct_height_count == num_rows:
            print(f"PASS: Component 2 — All {num_rows} rows have height ~{TARGET_TWIPS} twips (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 2 — {correct_height_count}/{num_rows} rows have correct height")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist unsaved edits before verification
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
