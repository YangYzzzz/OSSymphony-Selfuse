"""
Initial Setup: Legal brief appendix with 48 citation lines (16 duplicates among 32 unique)
Task ID: osworld_writer_duplicate_line_removal_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_duplicate_line_removal_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('APPENDIX A: TABLE OF AUTHORITIES', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    sub = doc.add_paragraph('Legal Citations — Consolidated Reference List')
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph('')  # blank line

    # 32 unique legal citations
    unique_citations = [
        "Anderson v. Liberty Lobby, Inc., 477 U.S. 242 (1986)",
        "Ashcroft v. Iqbal, 556 U.S. 662 (2009)",
        "Bell Atlantic Corp. v. Twombly, 550 U.S. 544 (2007)",
        "Brown v. Board of Education, 347 U.S. 483 (1954)",
        "Celotex Corp. v. Catrett, 477 U.S. 317 (1986)",
        "Citizens United v. Federal Election Commission, 558 U.S. 310 (2010)",
        "Dolan v. City of Tigard, 512 U.S. 374 (1994)",
        "Erie Railroad Co. v. Tompkins, 304 U.S. 64 (1938)",
        "Escobedo v. Illinois, 378 U.S. 478 (1964)",
        "Gideon v. Wainwright, 372 U.S. 335 (1963)",
        "Gonzales v. Raich, 545 U.S. 1 (2005)",
        "Griswold v. Connecticut, 381 U.S. 479 (1965)",
        "Hamdi v. Rumsfeld, 542 U.S. 507 (2004)",
        "Heart of Atlanta Motel v. United States, 379 U.S. 241 (1964)",
        "Heller v. District of Columbia, 554 U.S. 570 (2008)",
        "INS v. Chadha, 462 U.S. 919 (1983)",
        "Katz v. United States, 389 U.S. 347 (1967)",
        "Korematsu v. United States, 323 U.S. 214 (1944)",
        "Lemon v. Kurtzman, 403 U.S. 602 (1971)",
        "Mapp v. Ohio, 367 U.S. 643 (1961)",
        "Marbury v. Madison, 5 U.S. 137 (1803)",
        "McCulloch v. Maryland, 17 U.S. 316 (1819)",
        "Miranda v. Arizona, 384 U.S. 436 (1966)",
        "New York Times Co. v. Sullivan, 376 U.S. 254 (1964)",
        "Obergefell v. Hodges, 576 U.S. 644 (2015)",
        "Palko v. Connecticut, 302 U.S. 319 (1937)",
        "Plessy v. Ferguson, 163 U.S. 537 (1896)",
        "Roe v. Wade, 410 U.S. 113 (1973)",
        "Shelby County v. Holder, 570 U.S. 529 (2013)",
        "Terry v. Ohio, 392 U.S. 1 (1968)",
        "United States v. Nixon, 418 U.S. 683 (1974)",
        "Youngstown Sheet & Tube Co. v. Sawyer, 343 U.S. 579 (1952)",
    ]

    # 16 duplicates to scatter throughout the list (selecting from the 32 above)
    duplicates = [
        "Bell Atlantic Corp. v. Twombly, 550 U.S. 544 (2007)",
        "Brown v. Board of Education, 347 U.S. 483 (1954)",
        "Celotex Corp. v. Catrett, 477 U.S. 317 (1986)",
        "Erie Railroad Co. v. Tompkins, 304 U.S. 64 (1938)",
        "Gideon v. Wainwright, 372 U.S. 335 (1963)",
        "Hamdi v. Rumsfeld, 542 U.S. 507 (2004)",
        "Katz v. United States, 389 U.S. 347 (1967)",
        "Mapp v. Ohio, 367 U.S. 643 (1961)",
        "Marbury v. Madison, 5 U.S. 137 (1803)",
        "Miranda v. Arizona, 384 U.S. 436 (1966)",
        "New York Times Co. v. Sullivan, 376 U.S. 254 (1964)",
        "Obergefell v. Hodges, 576 U.S. 644 (2015)",
        "Roe v. Wade, 410 U.S. 113 (1973)",
        "Terry v. Ohio, 392 U.S. 1 (1968)",
        "United States v. Nixon, 418 U.S. 683 (1974)",
        "Youngstown Sheet & Tube Co. v. Sawyer, 343 U.S. 579 (1952)",
    ]

    # Build the 48-line list: interleave originals with duplicates at various positions
    # Original order (not alphabetical - scattered order to make sorting meaningful)
    original_order = [
        unique_citations[20],   # Marbury v. Madison
        unique_citations[2],    # Bell Atlantic (will be dup'd)
        unique_citations[7],    # Erie Railroad
        unique_citations[0],    # Anderson v. Liberty Lobby
        duplicates[8],          # Marbury v. Madison DUP
        unique_citations[22],   # Miranda v. Arizona
        unique_citations[10],   # Gonzales v. Raich
        duplicates[7],          # Mapp v. Ohio DUP
        unique_citations[3],    # Brown v. Board
        unique_citations[15],   # INS v. Chadha
        duplicates[1],          # Brown v. Board DUP
        unique_citations[19],   # Mapp v. Ohio
        unique_citations[27],   # Roe v. Wade
        duplicates[4],          # Gideon v. Wainwright DUP
        unique_citations[9],    # Gideon v. Wainwright
        unique_citations[24],   # Obergefell v. Hodges
        duplicates[11],         # Obergefell v. Hodges DUP
        unique_citations[4],    # Celotex Corp.
        unique_citations[1],    # Ashcroft v. Iqbal
        duplicates[2],          # Celotex Corp. DUP
        unique_citations[13],   # Heart of Atlanta
        unique_citations[31],   # Youngstown Sheet
        duplicates[15],         # Youngstown Sheet DUP
        unique_citations[17],   # Korematsu
        unique_citations[5],    # Citizens United
        duplicates[3],          # Erie Railroad DUP
        unique_citations[23],   # New York Times
        unique_citations[12],   # Hamdi v. Rumsfeld
        duplicates[5],          # Hamdi v. Rumsfeld DUP
        unique_citations[14],   # Heller v. DC
        unique_citations[8],    # Escobedo v. Illinois
        duplicates[9],          # Miranda v. Arizona DUP
        unique_citations[16],   # Katz v. United States
        unique_citations[6],    # Dolan v. City of Tigard
        duplicates[6],          # Katz v. United States DUP
        unique_citations[11],   # Griswold v. Connecticut
        unique_citations[18],   # Lemon v. Kurtzman
        duplicates[10],         # New York Times DUP
        unique_citations[21],   # McCulloch v. Maryland
        unique_citations[25],   # Palko v. Connecticut
        duplicates[12],         # Roe v. Wade DUP
        unique_citations[26],   # Plessy v. Ferguson
        unique_citations[28],   # Shelby County v. Holder
        duplicates[13],         # Terry v. Ohio DUP
        unique_citations[29],   # Terry v. Ohio
        unique_citations[30],   # United States v. Nixon
        duplicates[0],          # Bell Atlantic DUP
        duplicates[14],         # United States v. Nixon DUP
    ]

    # Verify we have exactly 48 entries
    assert len(original_order) == 48, f"Expected 48 citations, got {len(original_order)}"

    # Add section heading
    section_heading = doc.add_paragraph('CASES CITED')
    section_heading.runs[0].bold = True
    section_heading.runs[0].font.size = Pt(12)

    doc.add_paragraph('')  # blank line

    # Add all 48 citation lines as paragraphs
    for citation in original_order:
        para = doc.add_paragraph(citation)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.space_before = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Total citations: {len(original_order)} (32 unique + 16 duplicates)')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
