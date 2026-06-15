"""
Initial Setup: Create presentation notes document with tracked changes displayed
Task ID: writer_rm_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import lxml.etree as ET

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_normal_run(para, text, bold=False, italic=False, font_name="Calibri", font_size=11):
    """Add a normal (non-tracked) run to a paragraph."""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def add_tracked_insertion(para, text, author, date, rev_id):
    """Add a tracked insertion (w:ins) element to the paragraph."""
    ins = para._element.makeelement(qn('w:ins'), {
        qn('w:id'): str(rev_id),
        qn('w:author'): author,
        qn('w:date'): date,
    })
    r = ET.SubElement(ins, qn('w:r'))
    rpr = ET.SubElement(r, qn('w:rPr'))
    rn = ET.SubElement(rpr, qn('w:rFonts'), {qn('w:ascii'): 'Calibri', qn('w:hAnsi'): 'Calibri'})
    sz = ET.SubElement(rpr, qn('w:sz'), {qn('w:val'): '22'})  # 11pt
    t = ET.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    para._element.append(ins)
    return ins


def add_tracked_deletion(para, text, author, date, rev_id):
    """Add a tracked deletion (w:del) element to the paragraph."""
    de = para._element.makeelement(qn('w:del'), {
        qn('w:id'): str(rev_id),
        qn('w:author'): author,
        qn('w:date'): date,
    })
    r = ET.SubElement(de, qn('w:r'))
    rpr = ET.SubElement(r, qn('w:rPr'))
    rn = ET.SubElement(rpr, qn('w:rFonts'), {qn('w:ascii'): 'Calibri', qn('w:hAnsi'): 'Calibri'})
    sz = ET.SubElement(rpr, qn('w:sz'), {qn('w:val'): '22'})
    dt = ET.SubElement(r, qn('w:delText'))
    dt.text = text
    dt.set(qn('xml:space'), 'preserve')
    para._element.append(de)
    return de


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---- Title ----
    title = doc.add_heading('Q3 Product Launch Keynote - Speaker Notes', level=1)

    # ---- Section 1: Opening ----
    doc.add_heading('Opening Remarks (Slide 1-3)', level=2)

    p1 = doc.add_paragraph()
    add_normal_run(p1, 'Welcome the audience and thank them for attending the Q3 Product Launch event. ')
    add_normal_run(p1, 'Mention that this is our most ambitious release cycle to date, with ')
    # TRACKED CHANGE 1: Insertion - added "three" before "major"
    add_tracked_insertion(p1, 'three ', 'Maria Santos', '2026-03-28T09:15:00Z', 1)
    add_normal_run(p1, 'major feature updates across the platform.')

    p2 = doc.add_paragraph()
    add_normal_run(p2, 'Key talking points for the introduction:')

    # Bullet points
    b1 = doc.add_paragraph('Record quarterly revenue of $14.2M with 23% year-over-year growth', style='List Bullet')
    b2 = doc.add_paragraph(style='List Bullet')
    add_normal_run(b2, 'Customer base expanded to ')
    # TRACKED CHANGE 2: Deletion of old number, insertion of new
    add_tracked_deletion(b2, '8,500', 'Maria Santos', '2026-03-28T09:18:00Z', 2)
    add_tracked_insertion(b2, '12,400', 'Maria Santos', '2026-03-28T09:18:00Z', 3)
    add_normal_run(b2, ' active enterprise accounts')

    b3 = doc.add_paragraph('Strategic partnerships with Dataflow Inc. and CloudSync Technologies', style='List Bullet')

    # ---- Section 2: Product Demo ----
    doc.add_heading('Product Demo Walkthrough (Slide 4-8)', level=2)

    p3 = doc.add_paragraph()
    add_normal_run(p3, 'Begin the demo with the new dashboard redesign. Emphasize the ')
    add_normal_run(p3, 'drag-and-drop widget customization and real-time collaboration features. ')
    # TRACKED CHANGE 3: Insertion of a new sentence
    add_tracked_insertion(p3, 'Remember to show the mobile-responsive layout during this segment. ',
                          'David Chen', '2026-03-29T14:22:00Z', 4)
    add_normal_run(p3, 'Transition smoothly into the API integration demo.')

    p4 = doc.add_paragraph()
    add_normal_run(p4, 'Demo sequence:')

    doc.add_paragraph('Dashboard overview (2 minutes) - show filtering and sorting', style='List Number')
    doc.add_paragraph('Widget builder (3 minutes) - create a custom analytics panel', style='List Number')

    p_demo3 = doc.add_paragraph(style='List Number')
    # TRACKED CHANGE 4: Deletion of old timing
    add_tracked_deletion(p_demo3, 'API playground (2 minutes)', 'David Chen', '2026-03-29T14:25:00Z', 5)
    add_tracked_insertion(p_demo3, 'API playground (4 minutes)', 'David Chen', '2026-03-29T14:25:00Z', 6)
    add_normal_run(p_demo3, ' - live code examples with audience participation')

    doc.add_paragraph('Performance benchmarks (2 minutes) - side-by-side comparison', style='List Number')

    # ---- Section 3: Roadmap ----
    doc.add_heading('Future Roadmap (Slide 9-11)', level=2)

    p5 = doc.add_paragraph()
    add_normal_run(p5, 'Present the Q4 roadmap with confidence. Highlight the machine learning ')
    add_normal_run(p5, 'integration that the engineering team has been developing since January. ')
    add_normal_run(p5, 'Stress that all timelines are ')
    # TRACKED CHANGE 5: Deletion of "tentative" replaced with "confirmed"
    add_tracked_deletion(p5, 'tentative', 'Maria Santos', '2026-03-30T11:05:00Z', 7)
    add_tracked_insertion(p5, 'confirmed', 'Maria Santos', '2026-03-30T11:05:00Z', 8)
    add_normal_run(p5, ' and approved by the executive team.')

    p6 = doc.add_paragraph()
    add_normal_run(p6, 'Upcoming milestones:')

    doc.add_paragraph('ML-powered recommendations engine - Beta launch October 15', style='List Bullet')
    doc.add_paragraph('Enhanced security suite with SSO and RBAC - GA November 1', style='List Bullet')
    doc.add_paragraph('International expansion: EMEA data centers - December rollout', style='List Bullet')

    # ---- Section 4: Closing ----
    doc.add_heading('Closing & Q&A (Slide 12)', level=2)

    p7 = doc.add_paragraph()
    add_normal_run(p7, 'Summarize the three key takeaways: product innovation, customer growth, ')
    add_normal_run(p7, 'and strategic vision. Open the floor for questions. ')
    # TRACKED CHANGE 6: Insertion of reminder note
    add_tracked_insertion(p7, 'Prepare answers for pricing questions - refer to the updated pricing sheet shared by Finance on March 25.',
                          'David Chen', '2026-03-30T16:40:00Z', 9)

    p8 = doc.add_paragraph()
    add_normal_run(p8, 'Thank the audience and remind them about the networking reception ')
    add_normal_run(p8, 'in the Meridian Ballroom immediately following the presentation.')

    # ---- Enable "Show Changes" (tracked changes markup visible) ----
    # In OOXML, when there is NO <w:revisionView> element, or when
    # w:markup="1", tracked changes are shown. We ensure markup is ON
    # by either not adding revisionView or setting markup="1".
    # We'll add it explicitly with markup="1" for clarity.
    settings = doc.settings.element
    rv_existing = settings.findall(qn('w:revisionView'))
    for rv in rv_existing:
        settings.remove(rv)
    # Add revisionView with markup="1" (show changes ON)
    rv = settings.makeelement(qn('w:revisionView'), {
        qn('w:markup'): '1',
        qn('w:comments'): '1',
        qn('w:insDel'): '1',
        qn('w:formatting'): '1',
    })
    settings.append(rv)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
