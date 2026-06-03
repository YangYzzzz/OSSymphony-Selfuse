"""
Initial Setup: Create a legal contract document with a Definitions section
containing five defined terms in regular formatting.
Task ID: writer_legal_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(12)
    run = preamble.add_run(
        'This Professional Services Agreement (this "Agreement") is entered into '
        'as of March 15, 2025 (the "Effective Date"), by and between Nextera '
        'Solutions Inc., a Delaware corporation with its principal offices at '
        '2400 Technology Drive, Suite 800, San Jose, CA 95134 ("Provider"), and '
        'Meridian Healthcare Group LLC, a California limited liability company '
        'with its principal offices at 1750 Montgomery Street, San Francisco, '
        'CA 94111 ("Client").'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    recitals_heading = doc.add_heading('RECITALS', level=1)

    recital_a = doc.add_paragraph()
    run = recital_a.add_run(
        'WHEREAS, Provider is engaged in the business of providing information '
        'technology consulting, software development, and related professional '
        'services to enterprise clients; and'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    recital_b = doc.add_paragraph()
    run = recital_b.add_run(
        'WHEREAS, Client desires to engage Provider to perform certain '
        'professional services as more particularly described herein, and '
        'Provider desires to provide such services to Client, subject to the '
        'terms and conditions set forth in this Agreement.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    transition = doc.add_paragraph()
    run = transition.add_run(
        'NOW, THEREFORE, in consideration of the mutual covenants, promises, '
        'and agreements contained herein, and for other good and valuable '
        'consideration, the receipt and sufficiency of which are hereby '
        'acknowledged, the parties agree as follows:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 1: Definitions ---
    doc.add_heading('ARTICLE 1 \u2013 DEFINITIONS', level=1)

    intro = doc.add_paragraph()
    run = intro.add_run(
        'For purposes of this Agreement, the following terms shall have the '
        'meanings set forth below:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    definitions = [
        (
            '"Agreement"',
            ' means this Professional Services Agreement, including all '
            'exhibits, schedules, and amendments attached hereto or incorporated '
            'by reference herein, as may be amended from time to time in '
            'accordance with Section 12.3.'
        ),
        (
            '"Confidential Information"',
            ' means any and all non-public, proprietary, or confidential '
            'information disclosed by one party to the other party, whether '
            'orally, in writing, electronically, or by inspection of tangible '
            'objects, including but not limited to trade secrets, business '
            'plans, financial data, customer lists, technical specifications, '
            'software source code, and marketing strategies.'
        ),
        (
            '"Effective Date"',
            ' means the date first written above on which this Agreement '
            'becomes legally binding upon both parties, as specified in the '
            'preamble of this Agreement.'
        ),
        (
            '"Intellectual Property"',
            ' means all patents, copyrights, trademarks, service marks, trade '
            'secrets, know-how, inventions, discoveries, improvements, works of '
            'authorship, software, documentation, designs, and any other '
            'intellectual property rights, whether registered or unregistered, '
            'arising under the laws of any jurisdiction.'
        ),
        (
            '"Services"',
            ' means the professional consulting, software development, system '
            'integration, technical support, and related services to be '
            'performed by Provider for Client as described in Exhibit A '
            'attached hereto, and any additional services agreed upon by the '
            'parties in writing pursuant to a Statement of Work.'
        ),
    ]

    for term, definition in definitions:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.left_indent = Inches(0.5)

        # Term in quotes - regular formatting (NO bold, NO italic)
        run_term = para.add_run(term)
        run_term.font.size = Pt(11)
        run_term.font.name = 'Calibri'

        # Definition text
        run_def = para.add_run(definition)
        run_def.font.size = Pt(11)
        run_def.font.name = 'Calibri'

    # --- Section 2: Scope of Services ---
    doc.add_heading('ARTICLE 2 \u2013 SCOPE OF SERVICES', level=1)

    scope_text = doc.add_paragraph()
    run = scope_text.add_run(
        '2.1 Engagement. Client hereby engages Provider, and Provider hereby '
        'accepts such engagement, to perform the Services described in Exhibit A '
        'during the term of this Agreement. Provider shall perform the Services '
        'in a professional and workmanlike manner consistent with generally '
        'accepted industry standards and practices.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    scope_text2 = doc.add_paragraph()
    run = scope_text2.add_run(
        '2.2 Change Orders. Either party may request changes to the scope of '
        'Services by submitting a written change order to the other party. No '
        'change order shall be effective unless approved in writing by both '
        'parties and shall specify any adjustments to fees, timelines, or '
        'deliverables resulting from such change.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 3: Compensation ---
    doc.add_heading('ARTICLE 3 \u2013 COMPENSATION AND PAYMENT', level=1)

    comp_text = doc.add_paragraph()
    run = comp_text.add_run(
        '3.1 Fees. In consideration for the Services, Client shall pay Provider '
        'the fees set forth in Exhibit B (the "Fee Schedule"). Unless otherwise '
        'specified in a Statement of Work, Provider shall invoice Client monthly '
        'in arrears for Services rendered during the preceding calendar month.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    comp_text2 = doc.add_paragraph()
    run = comp_text2.add_run(
        '3.2 Payment Terms. Client shall pay all undisputed invoices within '
        'thirty (30) days of receipt. Late payments shall accrue interest at the '
        'rate of one and one-half percent (1.5%) per month or the maximum rate '
        'permitted by applicable law, whichever is less.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 4: Confidentiality ---
    doc.add_heading('ARTICLE 4 \u2013 CONFIDENTIALITY', level=1)

    conf_text = doc.add_paragraph()
    run = conf_text.add_run(
        '4.1 Non-Disclosure. Each party agrees to hold in strict confidence all '
        'Confidential Information received from the other party and shall not '
        'disclose such Confidential Information to any third party without the '
        'prior written consent of the disclosing party, except as required by '
        'law or court order.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    conf_text2 = doc.add_paragraph()
    run = conf_text2.add_run(
        '4.2 Return of Materials. Upon termination of this Agreement or upon '
        'request by the disclosing party, the receiving party shall promptly '
        'return or destroy all Confidential Information and certify in writing '
        'that it has done so.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 5: Intellectual Property ---
    doc.add_heading('ARTICLE 5 \u2013 INTELLECTUAL PROPERTY RIGHTS', level=1)

    ip_text = doc.add_paragraph()
    run = ip_text.add_run(
        '5.1 Ownership. All Intellectual Property created by Provider in the '
        'course of performing the Services shall be the exclusive property of '
        'Client upon full payment of all applicable fees. Provider hereby '
        'assigns to Client all right, title, and interest in and to such '
        'Intellectual Property.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    ip_text2 = doc.add_paragraph()
    run = ip_text2.add_run(
        '5.2 Pre-Existing IP. Notwithstanding the foregoing, Provider retains '
        'all rights in any pre-existing Intellectual Property that Provider '
        'incorporates into any deliverables. Provider grants Client a perpetual, '
        'non-exclusive, royalty-free license to use such pre-existing '
        'Intellectual Property solely in connection with the deliverables.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
