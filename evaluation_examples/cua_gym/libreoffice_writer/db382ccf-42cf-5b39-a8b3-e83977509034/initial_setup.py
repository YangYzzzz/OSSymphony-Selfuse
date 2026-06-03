"""
Initial Setup: Create User_Manual.docx with 22 pages and 8 key terms
Task ID: writer_pd_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_020'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body_text(doc, text):
    """Add a body paragraph with standard formatting."""
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    return p


def add_page_break(doc):
    """Add an explicit page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ========== PAGE 1: Title Page ==========
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('CloudSync Platform', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_heading('User Manual', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph('Version 3.2.1')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph('Prepared by: CloudSync Documentation Team')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph('Last Updated: March 2025')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph('Confidential — For Internal Use Only')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_page_break(doc)

    # ========== PAGE 2: Table of Contents ==========
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc_items = [
        ('1. Introduction', '3'),
        ('2. System Requirements', '4'),
        ('3. Installation Guide', '5'),
        ('4. Initial Configuration', '6'),
        ('5. User Authentication', '7'),
        ('6. Dashboard Overview', '8'),
        ('7. Data Management', '9'),
        ('8. Collaboration Tools', '10'),
        ('9. Reporting and Analytics', '11'),
        ('10. Notifications and Alerts', '12'),
        ('11. Advanced Settings', '13'),
        ('12. Backup and Recovery', '14'),
        ('13. Security Best Practices', '15'),
        ('14. Mobile Access', '16'),
        ('15. Troubleshooting Guide', '17'),
        ('16. Frequently Asked Questions', '18'),
        ('17. API Integration Guide', '19'),
        ('18. Release Notes', '20'),
        ('19. Glossary', '21'),
        ('20. Index', '22'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item} ')
        dots = '.' * (60 - len(item))
        p.add_run(dots)
        p.add_run(f' {page}')

    add_page_break(doc)

    # ========== PAGE 3: Introduction (contains "installation") ==========
    add_heading_styled(doc, '1. Introduction', level=1)
    add_body_text(doc, (
        'Welcome to the CloudSync Platform User Manual. This comprehensive guide covers '
        'everything you need to know about setting up and using CloudSync in your organization. '
        'Whether you are a new user or an experienced administrator, this manual provides '
        'detailed instructions for every aspect of the platform.'
    ))
    add_body_text(doc, (
        'CloudSync is an enterprise-grade cloud synchronization and collaboration platform '
        'designed for teams of all sizes. The platform supports real-time file synchronization, '
        'secure sharing, and comprehensive audit logging.'
    ))
    add_body_text(doc, (
        'Before you begin using CloudSync, you will need to complete the installation process '
        'on your local workstation. The installation procedure is straightforward and typically '
        'takes less than 10 minutes. Detailed steps for installation are provided in Chapter 3 '
        'of this manual.'
    ))
    add_body_text(doc, (
        'This manual is organized into 20 chapters, each covering a specific aspect of the '
        'platform. We recommend reading the chapters in order for first-time users, as each '
        'section builds upon concepts introduced in previous chapters.'
    ))
    add_body_text(doc, (
        'For technical support, please contact our help desk at support@cloudsync.example.com '
        'or call 1-800-CLOUDSYNC during business hours (Monday through Friday, 8 AM to 6 PM EST).'
    ))

    add_page_break(doc)

    # ========== PAGE 4: System Requirements ==========
    add_heading_styled(doc, '2. System Requirements', level=1)
    add_body_text(doc, (
        'Before proceeding with the setup, ensure your system meets the following minimum '
        'requirements for optimal performance with CloudSync Platform.'
    ))
    add_heading_styled(doc, 'Hardware Requirements', level=2)
    add_body_text(doc, 'Processor: Intel Core i5 (8th generation or later) or AMD Ryzen 5 equivalent')
    add_body_text(doc, 'Memory: 8 GB RAM minimum, 16 GB recommended for enterprise deployments')
    add_body_text(doc, 'Storage: 500 MB free disk space for application, plus storage for synced files')
    add_body_text(doc, 'Network: Broadband internet connection (minimum 10 Mbps upload/download)')

    add_heading_styled(doc, 'Software Requirements', level=2)
    add_body_text(doc, 'Operating System: Windows 10/11, macOS 12+, Ubuntu 20.04+, or RHEL 8+')
    add_body_text(doc, 'Browser: Chrome 90+, Firefox 88+, Safari 15+, or Edge 90+ for web interface')
    add_body_text(doc, 'Runtime: .NET 6.0 or later (bundled with installer)')
    add_body_text(doc, (
        'Additional dependencies are automatically resolved during the setup process. '
        'Enterprise customers may need to configure proxy settings prior to installation.'
    ))

    add_page_break(doc)

    # ========== PAGE 5: Installation Guide (contains "configuration") ==========
    add_heading_styled(doc, '3. Installation Guide', level=1)
    add_body_text(doc, (
        'This chapter walks you through the complete installation process for CloudSync '
        'on supported operating systems. Follow the steps carefully to ensure a successful setup.'
    ))
    add_heading_styled(doc, 'Downloading the Installer', level=2)
    add_body_text(doc, (
        'Visit https://downloads.cloudsync.example.com and select the appropriate installer '
        'for your operating system. Enterprise customers should use the MSI package for '
        'centralized deployment via Group Policy or SCCM.'
    ))
    add_heading_styled(doc, 'Running the Installer', level=2)
    add_body_text(doc, (
        'Double-click the downloaded installer file and follow the on-screen prompts. '
        'Accept the license agreement and choose your preferred installation directory. '
        'The default path is C:\\Program Files\\CloudSync on Windows or /opt/cloudsync on Linux.'
    ))
    add_body_text(doc, (
        'After the installation files are copied, the setup wizard will guide you through '
        'initial configuration steps. The configuration wizard covers network settings, '
        'proxy configuration, and initial account setup. Proper configuration at this stage '
        'ensures smooth operation going forward.'
    ))
    add_body_text(doc, (
        'Once the installer completes, CloudSync will appear in your system tray (Windows/Linux) '
        'or menu bar (macOS). Click the icon to open the sign-in dialog.'
    ))

    add_page_break(doc)

    # ========== PAGE 6: Initial Configuration ==========
    add_heading_styled(doc, '4. Initial Configuration', level=1)
    add_body_text(doc, (
        'After installation, you will need to configure CloudSync to connect to your '
        'organization\'s server. This chapter covers the essential settings you should '
        'adjust before first use.'
    ))
    add_heading_styled(doc, 'Server Connection', level=2)
    add_body_text(doc, (
        'Open CloudSync Settings and navigate to the Connection tab. Enter the server URL '
        'provided by your IT administrator (e.g., https://sync.yourcompany.com). Click '
        '"Test Connection" to verify connectivity.'
    ))
    add_heading_styled(doc, 'Sync Folder Selection', level=2)
    add_body_text(doc, (
        'Choose the local folder that will be synchronized with the cloud. By default, '
        'this is set to ~/CloudSync on macOS/Linux or C:\\Users\\<username>\\CloudSync on Windows. '
        'You can change this to any local directory with sufficient free space.'
    ))
    add_body_text(doc, (
        'Selective sync allows you to choose specific remote folders to sync locally, '
        'reducing bandwidth usage and disk space consumption. This is particularly useful '
        'for users on laptops with limited storage.'
    ))
    add_body_text(doc, (
        'Bandwidth throttling can be configured under Settings > Network to prevent CloudSync '
        'from consuming all available bandwidth during large synchronization operations.'
    ))

    add_page_break(doc)

    # ========== PAGE 7: User Authentication (contains "authentication") ==========
    add_heading_styled(doc, '5. User Authentication', level=1)
    add_body_text(doc, (
        'CloudSync supports multiple authentication methods to ensure secure access to your '
        'organization\'s data. This chapter describes the available authentication mechanisms '
        'and how to configure them.'
    ))
    add_heading_styled(doc, 'Single Sign-On (SSO)', level=2)
    add_body_text(doc, (
        'CloudSync integrates with SAML 2.0 and OpenID Connect identity providers. '
        'Administrators can configure SSO through the Admin Console under Security > '
        'Authentication. Supported providers include Okta, Azure AD, Google Workspace, '
        'and OneLogin.'
    ))
    add_heading_styled(doc, 'Two-Factor Authentication (2FA)', level=2)
    add_body_text(doc, (
        'For enhanced security, enable two-factor authentication from your profile settings. '
        'CloudSync supports TOTP-based authenticator apps (Google Authenticator, Authy) '
        'as well as hardware security keys (YubiKey, Titan). SMS-based 2FA is available '
        'but not recommended for high-security environments.'
    ))
    add_body_text(doc, (
        'Session management allows administrators to set maximum session duration, '
        'idle timeout periods, and enforce re-authentication for sensitive operations '
        'such as sharing external links or modifying security settings.'
    ))

    add_page_break(doc)

    # ========== PAGE 8: Dashboard Overview (contains "dashboard") ==========
    add_heading_styled(doc, '6. Dashboard Overview', level=1)
    add_body_text(doc, (
        'The CloudSync dashboard provides a centralized view of your synchronization status, '
        'recent activity, and quick access to frequently used features. Understanding the '
        'dashboard layout will help you navigate the platform efficiently.'
    ))
    add_heading_styled(doc, 'Main Dashboard Components', level=2)
    add_body_text(doc, (
        'The top section displays sync status indicators: a green checkmark means all files '
        'are synchronized, a blue rotating arrow indicates active syncing, and a red exclamation '
        'mark signals sync errors that require attention.'
    ))
    add_body_text(doc, (
        'The activity feed on the left shows recent file changes, sharing events, and system '
        'notifications in chronological order. Each entry includes the file name, action type, '
        'timestamp, and the user who performed the action.'
    ))
    add_heading_styled(doc, 'Quick Actions', level=2)
    add_body_text(doc, (
        'The right panel provides quick action buttons for common tasks: Upload Files, '
        'Create Folder, Share Link, and View Trash. These shortcuts save time compared '
        'to navigating through the full menu structure.'
    ))
    add_body_text(doc, (
        'You can customize the dashboard layout by dragging widgets to your preferred positions. '
        'Click the gear icon in the top-right corner to add or remove dashboard widgets.'
    ))

    add_page_break(doc)

    # ========== PAGE 9: Data Management ==========
    add_heading_styled(doc, '7. Data Management', level=1)
    add_body_text(doc, (
        'Effective data management is crucial for maintaining an organized and efficient '
        'workspace in CloudSync. This chapter covers file operations, version control, '
        'and storage management.'
    ))
    add_heading_styled(doc, 'File Operations', level=2)
    add_body_text(doc, (
        'CloudSync supports standard file operations including upload, download, rename, '
        'move, copy, and delete. All operations are logged in the audit trail for compliance '
        'purposes. Drag-and-drop is supported in the web interface for bulk uploads.'
    ))
    add_heading_styled(doc, 'Version History', level=2)
    add_body_text(doc, (
        'Every file modification creates a new version that is stored for 90 days by default. '
        'Administrators can adjust the retention period from 30 days to unlimited. To restore '
        'a previous version, right-click the file and select "Version History."'
    ))
    add_body_text(doc, (
        'CloudSync uses delta synchronization to minimize bandwidth usage. Only changed '
        'portions of files are transmitted, making large file updates significantly faster. '
        'This is especially beneficial for users working with large media files or databases.'
    ))

    add_page_break(doc)

    # ========== PAGE 10: Collaboration Tools ==========
    add_heading_styled(doc, '8. Collaboration Tools', level=1)
    add_body_text(doc, (
        'CloudSync provides robust collaboration features that enable teams to work together '
        'seamlessly, regardless of their physical location.'
    ))
    add_heading_styled(doc, 'Shared Folders', level=2)
    add_body_text(doc, (
        'Create shared folders to collaborate with team members. Permissions can be set at '
        'the folder or file level: View Only, Editor, or Admin. Shared folder members receive '
        'real-time notifications when files are added or modified.'
    ))
    add_heading_styled(doc, 'External Sharing', level=2)
    add_body_text(doc, (
        'Share files with external partners via secure links. Configure link expiration dates, '
        'download limits, and password protection. External recipients do not need a CloudSync '
        'account to access shared files.'
    ))
    add_heading_styled(doc, 'Comments and Annotations', level=2)
    add_body_text(doc, (
        'Add comments to files and folders to provide context or feedback. Use @mentions '
        'to notify specific team members. Comments support rich text formatting and file '
        'attachments for detailed discussions.'
    ))
    add_body_text(doc, (
        'Real-time co-editing is supported for document formats (.docx, .xlsx, .pptx) '
        'through the integrated LibreOffice Online editor. Changes are merged automatically '
        'with conflict resolution for simultaneous edits to the same section.'
    ))

    add_page_break(doc)

    # ========== PAGE 11: Reporting and Analytics (contains "reporting") ==========
    add_heading_styled(doc, '9. Reporting and Analytics', level=1)
    add_body_text(doc, (
        'CloudSync includes comprehensive reporting and analytics tools that help '
        'administrators monitor platform usage, track compliance, and optimize resource '
        'allocation. The reporting module provides both pre-built and custom report options.'
    ))
    add_heading_styled(doc, 'Standard Reports', level=2)
    add_body_text(doc, (
        'Pre-built reports include: Storage Usage by Department, Active User Summary, '
        'File Sharing Activity, Login Audit Trail, and Bandwidth Consumption. Reports can '
        'be exported in PDF, CSV, or Excel format.'
    ))
    add_heading_styled(doc, 'Custom Reports', level=2)
    add_body_text(doc, (
        'The report builder allows administrators to create custom reports by selecting '
        'data sources, filters, and visualization types. Saved reports can be scheduled '
        'for automatic generation and email delivery on a daily, weekly, or monthly basis.'
    ))
    add_body_text(doc, (
        'Analytics dashboards provide real-time insights into platform health, including '
        'sync performance metrics, error rates, and user engagement statistics. Threshold '
        'alerts can be configured to notify administrators when metrics exceed defined limits.'
    ))

    add_page_break(doc)

    # ========== PAGE 12: Notifications and Alerts ==========
    add_heading_styled(doc, '10. Notifications and Alerts', level=1)
    add_body_text(doc, (
        'Stay informed about important events with CloudSync\'s notification system. '
        'Notifications can be delivered via email, desktop push, mobile push, or in-app alerts.'
    ))
    add_heading_styled(doc, 'Notification Categories', level=2)
    add_body_text(doc, (
        'File Activity: New uploads, modifications, deletions, and sharing events. '
        'System Alerts: Sync errors, storage warnings, and maintenance windows. '
        'Security Events: Failed login attempts, permission changes, and suspicious activity.'
    ))
    add_heading_styled(doc, 'Customizing Notifications', level=2)
    add_body_text(doc, (
        'Navigate to Settings > Notifications to customize which events trigger alerts '
        'and through which channels. You can set quiet hours to suppress non-critical '
        'notifications during specified time periods.'
    ))
    add_body_text(doc, (
        'Administrators can configure organization-wide notification policies that enforce '
        'minimum alert levels for security events, ensuring critical notifications cannot '
        'be silenced by individual users.'
    ))

    add_page_break(doc)

    # ========== PAGE 13: Advanced Settings ==========
    add_heading_styled(doc, '11. Advanced Settings', level=1)
    add_body_text(doc, (
        'This chapter covers advanced configuration options for power users and administrators '
        'who need fine-grained control over CloudSync behavior.'
    ))
    add_heading_styled(doc, 'Sync Rules', level=2)
    add_body_text(doc, (
        'Create custom sync rules to exclude specific file types, folders, or file size '
        'thresholds. Rules use glob patterns (e.g., *.tmp, ~$*) and can be applied at the '
        'user, group, or organization level.'
    ))
    add_heading_styled(doc, 'Performance Tuning', level=2)
    add_body_text(doc, (
        'Adjust concurrent upload/download thread count (default: 4, maximum: 16). '
        'Configure chunk size for large file transfers (default: 8 MB). Enable or disable '
        'compression for specific file types. These settings affect both bandwidth usage '
        'and CPU load on client machines.'
    ))
    add_body_text(doc, (
        'Database maintenance tasks, including index optimization and log rotation, can be '
        'scheduled through the Admin Console under System > Maintenance. Regular maintenance '
        'helps prevent performance degradation over time.'
    ))

    add_page_break(doc)

    # ========== PAGE 14: Backup and Recovery (contains "backup") ==========
    add_heading_styled(doc, '12. Backup and Recovery', level=1)
    add_body_text(doc, (
        'Data protection is a core feature of CloudSync. This chapter describes the backup '
        'mechanisms available and the procedures for recovering lost or corrupted data. '
        'Understanding the backup system ensures your organization can respond effectively '
        'to data loss incidents.'
    ))
    add_heading_styled(doc, 'Automatic Backups', level=2)
    add_body_text(doc, (
        'CloudSync performs automatic backups of all synchronized data on a configurable '
        'schedule. The default backup frequency is every 6 hours with a 30-day retention '
        'period. Enterprise plans support continuous data protection (CDP) with point-in-time '
        'recovery capabilities.'
    ))
    add_heading_styled(doc, 'Manual Backup', level=2)
    add_body_text(doc, (
        'Administrators can trigger manual backups at any time from the Admin Console. '
        'Manual backups support full, incremental, and differential modes. Full backups '
        'capture the entire dataset while incremental backups only capture changes since '
        'the last backup, reducing storage requirements.'
    ))
    add_body_text(doc, (
        'Recovery procedures are initiated through the Admin Console under System > Recovery. '
        'Select the backup point, choose files or folders to restore, and specify the '
        'restoration target. Test restorations are recommended during quarterly DR drills.'
    ))

    add_page_break(doc)

    # ========== PAGE 15: Security Best Practices ==========
    add_heading_styled(doc, '13. Security Best Practices', level=1)
    add_body_text(doc, (
        'Maintaining strong security practices is essential for protecting sensitive data '
        'stored and synchronized through CloudSync. Follow these recommendations to minimize '
        'security risks.'
    ))
    add_heading_styled(doc, 'Access Control', level=2)
    add_body_text(doc, (
        'Apply the principle of least privilege when assigning permissions. Review shared '
        'folder access quarterly and revoke permissions for departed employees immediately. '
        'Use group-based permissions rather than individual assignments for easier management.'
    ))
    add_heading_styled(doc, 'Encryption', level=2)
    add_body_text(doc, (
        'CloudSync encrypts data at rest using AES-256 and in transit using TLS 1.3. '
        'Client-side encryption is available for highly sensitive files, ensuring that '
        'even CloudSync administrators cannot access the encrypted content. Encryption keys '
        'are managed through an integrated key management system or your organization\'s HSM.'
    ))
    add_body_text(doc, (
        'Regular security audits should be conducted using the built-in compliance scanner. '
        'The scanner checks for overshared files, expired external links, inactive accounts '
        'with elevated permissions, and other common security misconfigurations.'
    ))

    add_page_break(doc)

    # ========== PAGE 16: Mobile Access ==========
    add_heading_styled(doc, '14. Mobile Access', level=1)
    add_body_text(doc, (
        'Access your CloudSync files on the go with our mobile applications for iOS and '
        'Android. The mobile apps provide a streamlined interface optimized for touch '
        'interaction while maintaining the security features of the desktop client.'
    ))
    add_heading_styled(doc, 'Mobile App Features', level=2)
    add_body_text(doc, (
        'Browse and search files, preview documents and images, share files via links, '
        'upload photos and videos directly from your camera roll, and receive push '
        'notifications for file activity. Offline mode allows you to mark files for '
        'offline access on your device.'
    ))
    add_heading_styled(doc, 'Mobile Security', level=2)
    add_body_text(doc, (
        'The mobile app supports biometric authentication (fingerprint, Face ID), '
        'app-level PIN protection, and remote wipe capabilities. MDM integration with '
        'Microsoft Intune, VMware Workspace ONE, and Jamf is supported for enterprise '
        'device management.'
    ))
    add_body_text(doc, (
        'Data cached on mobile devices is encrypted using the device\'s secure enclave. '
        'Administrators can set policies to restrict file downloads to managed devices only '
        'and enforce minimum OS version requirements.'
    ))

    add_page_break(doc)

    # ========== PAGE 17: Troubleshooting Guide (contains "troubleshooting") ==========
    add_heading_styled(doc, '15. Troubleshooting Guide', level=1)
    add_body_text(doc, (
        'This chapter provides solutions for common issues encountered when using CloudSync. '
        'If your issue is not covered here, contact our support team for assistance. '
        'Effective troubleshooting often begins with checking the sync log files located '
        'in the application data directory.'
    ))
    add_heading_styled(doc, 'Sync Issues', level=2)
    add_body_text(doc, (
        'Problem: Files not syncing. Solution: Check your internet connection, verify the '
        'CloudSync service is running, ensure the file is not locked by another application, '
        'and confirm you have sufficient storage quota remaining.'
    ))
    add_body_text(doc, (
        'Problem: Conflicted copies appearing. Solution: Conflicted copies are created when '
        'the same file is modified simultaneously on multiple devices. Open both versions, '
        'merge changes manually, and delete the conflicted copy.'
    ))
    add_heading_styled(doc, 'Performance Issues', level=2)
    add_body_text(doc, (
        'Problem: Slow synchronization. Solution: Check bandwidth throttling settings, '
        'reduce the number of concurrent transfers, exclude large temporary files from sync, '
        'and verify your network latency to the CloudSync server is within acceptable limits.'
    ))
    add_body_text(doc, (
        'Problem: High CPU usage. Solution: Reduce the file system monitoring frequency, '
        'increase the minimum sync interval, and consider excluding directories with '
        'frequently changing temporary files.'
    ))

    add_page_break(doc)

    # ========== PAGE 18: FAQ ==========
    add_heading_styled(doc, '16. Frequently Asked Questions', level=1)
    add_body_text(doc, (
        'Q: How much storage do I get with CloudSync?\n'
        'A: Individual plans include 100 GB. Business plans start at 1 TB per user with '
        'unlimited storage available on Enterprise plans.'
    ))
    add_body_text(doc, (
        'Q: Can I share files with people outside my organization?\n'
        'A: Yes. Use the external sharing feature to create secure links. Recipients do not '
        'need a CloudSync account. Links can be password-protected and time-limited.'
    ))
    add_body_text(doc, (
        'Q: What happens if I accidentally delete a file?\n'
        'A: Deleted files are moved to the Trash and retained for 30 days. Administrators '
        'can restore files from the Trash or from backups beyond the retention period.'
    ))
    add_body_text(doc, (
        'Q: Is CloudSync HIPAA compliant?\n'
        'A: Yes. CloudSync Enterprise plans include HIPAA compliance features including '
        'BAA signing, audit logging, access controls, and encryption at rest. Contact our '
        'compliance team for a detailed compliance matrix.'
    ))
    add_body_text(doc, (
        'Q: How do I migrate data from another cloud storage provider?\n'
        'A: Use the CloudSync Migration Tool available from the Admin Console. It supports '
        'direct migration from Google Drive, Dropbox, OneDrive, and Box with folder structure '
        'and sharing permission preservation.'
    ))

    add_page_break(doc)

    # ========== PAGE 19: API Integration Guide (contains "API integration") ==========
    add_heading_styled(doc, '17. API Integration Guide', level=1)
    add_body_text(doc, (
        'CloudSync provides a comprehensive RESTful API for programmatic access to all '
        'platform features. This chapter introduces the API integration capabilities and '
        'provides guidance for developers building custom integrations.'
    ))
    add_heading_styled(doc, 'API Authentication', level=2)
    add_body_text(doc, (
        'All API requests require authentication via OAuth 2.0 bearer tokens. Register your '
        'application in the Developer Portal to obtain client credentials. The API supports '
        'both authorization code flow (for user-facing apps) and client credentials flow '
        '(for server-to-server integration).'
    ))
    add_heading_styled(doc, 'Core Endpoints', level=2)
    add_body_text(doc, (
        'GET /api/v2/files — List files and folders. POST /api/v2/files/upload — Upload '
        'a new file. PUT /api/v2/files/{id} — Update file metadata. DELETE /api/v2/files/{id} '
        '— Move file to trash. GET /api/v2/users — List organization users.'
    ))
    add_body_text(doc, (
        'The API integration supports webhook notifications for real-time event streaming. '
        'Configure webhooks in the Developer Portal to receive HTTP callbacks for file changes, '
        'user events, and administrative actions. Rate limiting is applied at 1000 requests '
        'per minute per API key.'
    ))
    add_body_text(doc, (
        'SDKs are available for Python, JavaScript, Java, Go, and C#. Each SDK includes '
        'comprehensive documentation, code samples, and integration test suites. Visit '
        'https://developers.cloudsync.example.com for the complete API reference.'
    ))

    add_page_break(doc)

    # ========== PAGE 20: Release Notes ==========
    add_heading_styled(doc, '18. Release Notes', level=1)
    add_heading_styled(doc, 'Version 3.2.1 (March 2025)', level=2)
    add_body_text(doc, (
        'Bug Fixes: Resolved intermittent sync failures on macOS Sequoia. Fixed memory leak '
        'in file watcher service on Windows. Corrected time zone handling in audit log exports.'
    ))
    add_heading_styled(doc, 'Version 3.2.0 (January 2025)', level=2)
    add_body_text(doc, (
        'New Features: Real-time co-editing for spreadsheets. Enhanced search with full-text '
        'indexing. New admin dashboard widgets for storage analytics. Improved mobile app '
        'performance with lazy loading.'
    ))
    add_heading_styled(doc, 'Version 3.1.0 (October 2024)', level=2)
    add_body_text(doc, (
        'New Features: Client-side encryption, SCIM provisioning support, custom branding '
        'for enterprise portals. Improvements: 40% faster delta sync algorithm, reduced '
        'memory footprint on Linux clients.'
    ))
    add_body_text(doc, (
        'For a complete list of changes in all versions, visit the CloudSync Release History '
        'page at https://docs.cloudsync.example.com/releases.'
    ))

    add_page_break(doc)

    # ========== PAGE 21: Glossary ==========
    add_heading_styled(doc, '19. Glossary', level=1)
    glossary_terms = [
        ('CDP', 'Continuous Data Protection — a method of backing up data by automatically saving changes in real time.'),
        ('Delta Sync', 'A synchronization method that transfers only the changed portions of a file rather than the entire file.'),
        ('HSM', 'Hardware Security Module — a physical device for managing digital keys and performing cryptographic operations.'),
        ('MDM', 'Mobile Device Management — software used to manage and secure mobile devices in enterprise environments.'),
        ('OAuth 2.0', 'An open authorization framework that provides secure delegated access to server resources.'),
        ('SAML', 'Security Assertion Markup Language — an XML-based standard for exchanging authentication data between parties.'),
        ('SCIM', 'System for Cross-domain Identity Management — a protocol for automating user provisioning and deprovisioning.'),
        ('SSO', 'Single Sign-On — an authentication scheme allowing users to log in once and access multiple applications.'),
        ('TLS', 'Transport Layer Security — a cryptographic protocol that provides secure communication over a network.'),
        ('TOTP', 'Time-based One-Time Password — an algorithm for generating temporary authentication codes.'),
    ]
    for term, definition in glossary_terms:
        p = doc.add_paragraph()
        run = p.add_run(f'{term}: ')
        run.bold = True
        p.add_run(definition)

    add_page_break(doc)

    # ========== PAGE 22: Index (empty - task is to populate this) ==========
    add_heading_styled(doc, '20. Index', level=1)
    add_body_text(doc, '')  # Empty content below heading — task is to generate the index here

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
