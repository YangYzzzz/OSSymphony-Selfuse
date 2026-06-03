"""
Initial Setup: Raw training content in plain text for 5 onboarding modules
Task ID: writer_hr_067
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_067'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Raw training content, plain text only ---
    # No title page, no TOC, no outline numbering, no callout boxes,
    # no captioned frames, no glossary, no appendix, no page numbers.

    # Title area (just plain text, no special formatting)
    doc.add_paragraph("Nextera Solutions Inc.")
    doc.add_paragraph("Employee Onboarding Training Manual")
    doc.add_paragraph("Human Resources Department")
    doc.add_paragraph("Effective Date: March 15, 2026")
    doc.add_paragraph("Version 2.4")
    doc.add_paragraph("")

    # Module 1
    doc.add_paragraph("Module 1: Company Overview and Culture")
    doc.add_paragraph("")
    doc.add_paragraph("Welcome to Nextera Solutions")
    doc.add_paragraph(
        "Welcome to Nextera Solutions Inc., a global leader in enterprise software and cloud "
        "infrastructure. Founded in 2008 by Dr. Amelia Thornton and Rajesh Kapoor, the company "
        "has grown from a 12-person startup in Austin, Texas to a multinational organization "
        "with over 4,500 employees across 18 offices worldwide."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Our Mission and Values")
    doc.add_paragraph(
        "Our mission is to empower businesses through innovative technology solutions that "
        "drive efficiency, scalability, and sustainable growth. We are guided by five core values:"
    )
    doc.add_paragraph("Innovation First - We challenge conventional thinking and embrace creative problem-solving.")
    doc.add_paragraph("Customer Obsession - Every decision starts with how it impacts our clients.")
    doc.add_paragraph("Integrity Always - We maintain the highest ethical standards in all interactions.")
    doc.add_paragraph("Collaborative Spirit - Great achievements come from diverse teams working together.")
    doc.add_paragraph("Continuous Learning - We invest in personal and professional development at every level.")
    doc.add_paragraph("")
    doc.add_paragraph("Organizational Structure")
    doc.add_paragraph(
        "Nextera Solutions operates through four primary divisions: Enterprise Products, "
        "Cloud Services, Professional Services, and Research & Development. Each division "
        "is led by a Senior Vice President who reports directly to the CEO, Dr. Thornton."
    )
    doc.add_paragraph("")

    # Safety warning content (plain text, no callout box)
    doc.add_paragraph(
        "Important: All employees must complete the Workplace Safety Orientation within their "
        "first 48 hours. Failure to complete this requirement may result in restricted building "
        "access until the training is finished."
    )
    doc.add_paragraph("")

    # Module 2
    doc.add_paragraph("Module 2: IT Systems and Security Protocols")
    doc.add_paragraph("")
    doc.add_paragraph("Account Setup Procedures")
    doc.add_paragraph(
        "Your IT onboarding begins with the provisioning of essential accounts and systems. "
        "The following steps must be completed in order:"
    )
    doc.add_paragraph("Contact IT Help Desk at extension 4500 or helpdesk@nextera.com to initiate account creation.")
    doc.add_paragraph("Provide your employee ID number from your offer letter.")
    doc.add_paragraph("IT will provision your Active Directory account within 2 business hours.")
    doc.add_paragraph("You will receive credentials for email (Outlook), Slack, Jira, and Confluence.")
    doc.add_paragraph("Set up multi-factor authentication (MFA) using the Authenticator app on your mobile device.")
    doc.add_paragraph("Complete the VPN configuration for remote access capability.")
    doc.add_paragraph("")
    doc.add_paragraph("Network Security Requirements")
    doc.add_paragraph(
        "All employees must adhere to the following security protocols to protect company "
        "assets and client data:"
    )
    doc.add_paragraph("Passwords must be at least 14 characters, including uppercase, lowercase, numbers, and symbols.")
    doc.add_paragraph("Passwords expire every 90 days and cannot repeat any of the last 12 passwords.")
    doc.add_paragraph("Screen lock must activate after 5 minutes of inactivity.")
    doc.add_paragraph("USB storage devices are prohibited on company workstations.")
    doc.add_paragraph("All software installations require IT approval via the Software Request Portal.")
    doc.add_paragraph("")

    # Screenshot placeholder text (plain text, no frame)
    doc.add_paragraph("[Screenshot: VPN Configuration Dashboard - shows the setup wizard interface]")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Important: Never share your credentials with anyone, including IT staff. Nextera IT "
        "will never ask for your password. Report any suspicious requests to security@nextera.com "
        "immediately."
    )
    doc.add_paragraph("")

    # Module 3
    doc.add_paragraph("Module 3: Benefits and Compensation")
    doc.add_paragraph("")
    doc.add_paragraph("Health and Wellness Benefits")
    doc.add_paragraph(
        "Nextera Solutions provides a comprehensive benefits package that begins on your first "
        "day of employment. You must complete benefits enrollment within 30 days of your start "
        "date through the HR Portal."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Enrollment Steps")
    doc.add_paragraph("Log into the HR Portal at hr.nextera.com using your company credentials.")
    doc.add_paragraph("Navigate to Benefits Enrollment under the My Benefits tab.")
    doc.add_paragraph("Review available plans including Medical, Dental, Vision, and Life Insurance.")
    doc.add_paragraph("Select your coverage tier: Employee Only, Employee + Spouse, Employee + Children, or Family.")
    doc.add_paragraph("Designate beneficiaries for life insurance and retirement accounts.")
    doc.add_paragraph("Submit enrollment and download confirmation PDF for your records.")
    doc.add_paragraph("")
    doc.add_paragraph("Retirement and Financial Planning")
    doc.add_paragraph(
        "The company offers a 401(k) plan with employer matching up to 6% of your annual salary. "
        "Vesting occurs on a graduated schedule: 20% after year one, 40% after year two, 60% "
        "after year three, 80% after year four, and 100% after year five."
    )
    doc.add_paragraph("")

    doc.add_paragraph("[Screenshot: HR Portal Benefits Dashboard - enrollment interface overview]")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Important: Benefits elections are irrevocable until the next open enrollment period "
        "in November, unless you experience a qualifying life event such as marriage, birth of "
        "a child, or loss of other coverage."
    )
    doc.add_paragraph("")

    # Module 4
    doc.add_paragraph("Module 4: Workplace Policies and Compliance")
    doc.add_paragraph("")
    doc.add_paragraph("Code of Conduct")
    doc.add_paragraph(
        "All employees are expected to maintain the highest standards of professional behavior. "
        "The Nextera Code of Conduct applies to all business activities, communications, and "
        "interactions with colleagues, clients, and partners."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Anti-Harassment Policy")
    doc.add_paragraph(
        "Nextera Solutions maintains a zero-tolerance policy for harassment of any kind, including "
        "but not limited to: verbal, physical, visual, or online harassment based on race, gender, "
        "age, religion, disability, sexual orientation, or any other protected characteristic."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Reporting Procedures")
    doc.add_paragraph("Speak directly with your manager if comfortable doing so.")
    doc.add_paragraph("Contact the HR Business Partner assigned to your department.")
    doc.add_paragraph("Use the anonymous Ethics Hotline at 1-888-555-0142.")
    doc.add_paragraph("Submit a report through the online Ethics Portal at ethics.nextera.com.")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Important: Retaliation against any employee who reports a concern in good faith is "
        "strictly prohibited and will result in disciplinary action up to and including termination."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Data Privacy and Confidentiality")
    doc.add_paragraph(
        "Employees handling client data must complete the Data Privacy Certification course "
        "within 60 days of hire. This certification must be renewed annually. Violations of "
        "data privacy policies may result in termination and legal action."
    )
    doc.add_paragraph("")

    doc.add_paragraph("[Screenshot: Ethics Portal Reporting Interface - anonymous submission form]")
    doc.add_paragraph("")

    # Module 5
    doc.add_paragraph("Module 5: Professional Development and Career Growth")
    doc.add_paragraph("")
    doc.add_paragraph("Learning and Development Programs")
    doc.add_paragraph(
        "Nextera Solutions invests heavily in employee growth through structured learning "
        "programs. Each employee receives an annual learning budget of $3,500 for conferences, "
        "courses, certifications, and educational materials."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Available Programs")
    doc.add_paragraph("Technical Skills Track - Coding bootcamps, cloud certifications (AWS, Azure, GCP), and architecture workshops.")
    doc.add_paragraph("Leadership Development - Management training for aspiring and current team leads.")
    doc.add_paragraph("Cross-Functional Rotation - 6-month rotations to build breadth across divisions.")
    doc.add_paragraph("Mentorship Program - Pairing with senior leaders for career guidance and networking.")
    doc.add_paragraph("Tuition Reimbursement - Up to $10,000 per year for degree programs related to your role.")
    doc.add_paragraph("")
    doc.add_paragraph("Performance Review Process")
    doc.add_paragraph(
        "Performance reviews are conducted semi-annually in June and December. The review "
        "process includes self-assessment, peer feedback (360-degree), manager evaluation, "
        "and a calibration session. Performance ratings use a 5-point scale:"
    )
    doc.add_paragraph("Exceptional (5) - Consistently exceeds expectations and demonstrates outstanding contributions.")
    doc.add_paragraph("Exceeds Expectations (4) - Regularly surpasses performance standards.")
    doc.add_paragraph("Meets Expectations (3) - Consistently meets all job requirements and standards.")
    doc.add_paragraph("Needs Improvement (2) - Performance falls short in some areas; development plan required.")
    doc.add_paragraph("Unsatisfactory (1) - Fails to meet minimum standards; formal performance improvement plan initiated.")
    doc.add_paragraph("")

    doc.add_paragraph("[Screenshot: Learning Management System - course catalog and enrollment page]")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Important: Employees must complete a minimum of 40 hours of professional development "
        "per year to remain in good standing. Tracking is done through the Learning Management "
        "System (LMS) at learn.nextera.com."
    )
    doc.add_paragraph("")

    # Glossary content (plain text, no special section formatting)
    doc.add_paragraph("Glossary")
    doc.add_paragraph("")
    doc.add_paragraph("Active Directory (AD) - Microsoft's directory service for managing network resources and user authentication.")
    doc.add_paragraph("MFA (Multi-Factor Authentication) - Security method requiring two or more verification factors to access an account.")
    doc.add_paragraph("VPN (Virtual Private Network) - Encrypted connection allowing secure remote access to company network.")
    doc.add_paragraph("401(k) - Employer-sponsored retirement savings plan with tax advantages.")
    doc.add_paragraph("HRIS (Human Resource Information System) - Software platform for managing HR data and processes.")
    doc.add_paragraph("PIP (Performance Improvement Plan) - Structured plan to help employees address performance deficiencies.")
    doc.add_paragraph("SLA (Service Level Agreement) - Contractual commitment defining expected service standards and response times.")
    doc.add_paragraph("GDPR (General Data Protection Regulation) - European Union regulation governing data privacy and protection.")
    doc.add_paragraph("SOC 2 (Service Organization Control 2) - Auditing standard for service providers managing customer data.")
    doc.add_paragraph("CI/CD (Continuous Integration/Continuous Deployment) - Software development practice automating testing and deployment.")
    doc.add_paragraph("")

    # Appendix content (plain text, no tables)
    doc.add_paragraph("Appendix A: Key Contacts and Resources")
    doc.add_paragraph("")
    doc.add_paragraph("Department - Contact - Phone - Email")
    doc.add_paragraph("IT Help Desk - Service Desk Team - Ext. 4500 - helpdesk@nextera.com")
    doc.add_paragraph("Human Resources - Maria Santos - Ext. 3200 - maria.santos@nextera.com")
    doc.add_paragraph("Facilities - Tom Henderson - Ext. 2800 - tom.henderson@nextera.com")
    doc.add_paragraph("Security - James Park - Ext. 5100 - security@nextera.com")
    doc.add_paragraph("Benefits Admin - Linda Zhao - Ext. 3215 - benefits@nextera.com")
    doc.add_paragraph("Payroll - David Kim - Ext. 3220 - payroll@nextera.com")
    doc.add_paragraph("Legal & Compliance - Sarah Mitchell - Ext. 6000 - compliance@nextera.com")
    doc.add_paragraph("Learning & Dev - Priya Sharma - Ext. 3250 - learning@nextera.com")
    doc.add_paragraph("")

    doc.add_paragraph("Appendix B: Training Completion Checklist")
    doc.add_paragraph("")
    doc.add_paragraph("Module - Requirement - Deadline - Status")
    doc.add_paragraph("Module 1 - Company Overview - Day 1 - Pending")
    doc.add_paragraph("Module 2 - IT Systems Setup - Day 2 - Pending")
    doc.add_paragraph("Module 3 - Benefits Enrollment - Day 5 - Pending")
    doc.add_paragraph("Module 4 - Compliance Training - Day 10 - Pending")
    doc.add_paragraph("Module 5 - Development Plan - Day 30 - Pending")
    doc.add_paragraph("Safety Orientation - In-Person Session - 48 Hours - Pending")
    doc.add_paragraph("Data Privacy Cert - Online Course - 60 Days - Pending")
    doc.add_paragraph("Ethics Acknowledgment - Signed Form - Day 5 - Pending")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
