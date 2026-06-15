"""
Reward Script: Comprehensive training manual with advanced formatting
Task ID: writer_hr_067
Domain: libreoffice_writer
Scoring:
  C1 (0.15) - Heading styles (Heading 1/2/3) applied
  C2 (0.15) - Outline numbering (1, 1.1, 1.1.1 format)
  C3 (0.10) - Table of Contents with TOC field code
  C4 (0.15) - 'Important' callout boxes with colored shading/borders
  C5 (0.10) - Screenshot placeholder frames with captions
  C6 (0.10) - Glossary section with term/definition table
  C7 (0.10) - Appendix section with reference tables
  C8 (0.10) - Page numbers starting after TOC (multi-section + PAGE field)
  C9 (0.05) - Title page formatting (centered, bold, large font)
"""

import os
import re

from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_067'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- Component 1: Heading styles applied (0.15 points) ----------
    # The initial file has NO headings (all Normal). Golden has Heading 1/2/3.
    try:
        h1_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 1')
        h2_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 2')
        h3_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 3')
        total_headings = h1_count + h2_count + h3_count

        # Need all three levels for full credit, with at least 5 H1, 5 H2, 3 H3
        if h1_count >= 5 and h2_count >= 5 and h3_count >= 3:
            print(f"PASS: Component 1 — Heading styles: H1={h1_count}, H2={h2_count}, H3={h3_count} (0.15 pts)")
            total_score += 0.15
        elif total_headings >= 5 and h1_count >= 2:
            print(f"PARTIAL: Component 1 — Some headings found: H1={h1_count}, H2={h2_count}, H3={h3_count} (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 1 — Insufficient headings: H1={h1_count}, H2={h2_count}, H3={h3_count}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------- Component 2: Outline numbering in headings (0.15 points) ----------
    # Golden has 1, 1.1, 1.1.1 numbering pattern in heading text
    try:
        numbered_headings = 0
        has_level1 = False
        has_level2 = False
        has_level3 = False

        for p in doc.paragraphs:
            if p.style.name.startswith('Heading'):
                match = re.match(r'^(\d+(\.\d+)*)\s', p.text)
                if match:
                    numbered_headings += 1
                    num = match.group(1)
                    parts = num.split('.')
                    if len(parts) == 1:
                        has_level1 = True
                    elif len(parts) == 2:
                        has_level2 = True
                    elif len(parts) == 3:
                        has_level3 = True

        if has_level1 and has_level2 and has_level3 and numbered_headings >= 10:
            print(f"PASS: Component 2 — Outline numbering: {numbered_headings} numbered headings with 3 levels (0.15 pts)")
            total_score += 0.15
        elif has_level1 and has_level2 and numbered_headings >= 5:
            print(f"PARTIAL: Component 2 — Partial outline numbering: {numbered_headings} numbered, levels 1&2 (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 2 — Outline numbering insufficient: {numbered_headings} numbered, L1={has_level1}, L2={has_level2}, L3={has_level3}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------- Component 3: Table of Contents with TOC field (0.10 points) ----------
    # Golden has a "Table of Contents" heading and a TOC field code
    try:
        has_toc_heading = False
        has_toc_field = False

        for p in doc.paragraphs:
            if 'table of contents' in p.text.lower() and p.style.name.startswith('Heading'):
                has_toc_heading = True
            # Check for TOC field code in paragraph XML
            for instr in p._element.findall('.//' + qn('w:instrText')):
                if instr.text and 'TOC' in instr.text:
                    has_toc_field = True

        if has_toc_heading and has_toc_field:
            print(f"PASS: Component 3 — TOC heading and TOC field code found (0.10 pts)")
            total_score += 0.10
        elif has_toc_heading or has_toc_field:
            print(f"PARTIAL: Component 3 — TOC heading={has_toc_heading}, TOC field={has_toc_field} (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — No TOC heading or field found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------- Component 4: Important callout boxes with colored borders (0.15 points) ----------
    # Golden has single-cell tables with "Important" text and colored shading (FFF2F2)
    try:
        callout_count = 0
        callout_with_shading = 0

        for t in doc.tables:
            if len(t.rows) == 1 and len(t.columns) == 1:
                cell_text = t.cell(0, 0).text.strip()
                if 'important' in cell_text.lower():
                    callout_count += 1
                    # Check for colored shading/borders
                    tc = t.cell(0, 0)._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        borders = tcPr.find(qn('w:tcBorders'))
                        if shd is not None or borders is not None:
                            callout_with_shading += 1

        if callout_with_shading >= 4:
            print(f"PASS: Component 4 — {callout_with_shading} callout boxes with colored borders/shading (0.15 pts)")
            total_score += 0.15
        elif callout_with_shading >= 2:
            print(f"PARTIAL: Component 4 — {callout_with_shading} callout boxes with styling (0.08 pts)")
            total_score += 0.08
        elif callout_count >= 2:
            print(f"PARTIAL: Component 4 — {callout_count} callout boxes but no colored styling (0.04 pts)")
            total_score += 0.04
        else:
            print(f"FAIL: Component 4 — Found {callout_count} callout boxes, {callout_with_shading} with styling")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---------- Component 5: Screenshot placeholders with captions (0.10 points) ----------
    # Golden has tables with "[Screenshot Placeholder]" text, and Figure X.X captions nearby
    try:
        placeholder_count = 0
        figure_caption_count = 0

        for t in doc.tables:
            cell_text = t.cell(0, 0).text.strip()
            if 'screenshot placeholder' in cell_text.lower() or 'placeholder' in cell_text.lower():
                placeholder_count += 1

        for p in doc.paragraphs:
            if re.match(r'^Figure\s+\d+\.\d+:', p.text):
                figure_caption_count += 1

        if placeholder_count >= 3 and figure_caption_count >= 3:
            print(f"PASS: Component 5 — {placeholder_count} placeholders, {figure_caption_count} figure captions (0.10 pts)")
            total_score += 0.10
        elif placeholder_count >= 2 or figure_caption_count >= 2:
            print(f"PARTIAL: Component 5 — {placeholder_count} placeholders, {figure_caption_count} captions (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — placeholders={placeholder_count}, captions={figure_caption_count}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ---------- Component 6: Glossary section with term/definition table (0.10 points) ----------
    # Golden has "Glossary" Heading 1 and a table with Term/Definition columns
    try:
        has_glossary_heading = False
        has_glossary_table = False

        for p in doc.paragraphs:
            if 'glossary' in p.text.lower() and p.style.name.startswith('Heading'):
                has_glossary_heading = True

        for t in doc.tables:
            if len(t.rows) >= 5 and len(t.columns) == 2:
                header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
                if 'term' in header_cells and 'definition' in header_cells:
                    has_glossary_table = True

        if has_glossary_heading and has_glossary_table:
            print(f"PASS: Component 6 — Glossary heading and term/definition table found (0.10 pts)")
            total_score += 0.10
        elif has_glossary_heading:
            print(f"PARTIAL: Component 6 — Glossary heading found but no proper table (0.04 pts)")
            total_score += 0.04
        else:
            print(f"FAIL: Component 6 — glossary_heading={has_glossary_heading}, glossary_table={has_glossary_table}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # ---------- Component 7: Appendix with reference tables (0.10 points) ----------
    # Golden has Appendix A and B headings with reference tables
    try:
        appendix_headings = 0
        appendix_tables = 0

        for p in doc.paragraphs:
            if 'appendix' in p.text.lower() and p.style.name.startswith('Heading'):
                appendix_headings += 1

        # Check for multi-column, multi-row tables (the reference tables in appendix)
        for t in doc.tables:
            if len(t.rows) >= 5 and len(t.columns) >= 3:
                appendix_tables += 1

        if appendix_headings >= 2 and appendix_tables >= 1:
            print(f"PASS: Component 7 — {appendix_headings} appendix headings, {appendix_tables} reference tables (0.10 pts)")
            total_score += 0.10
        elif appendix_headings >= 1:
            print(f"PARTIAL: Component 7 — {appendix_headings} appendix headings, {appendix_tables} tables (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 — appendix_headings={appendix_headings}, reference_tables={appendix_tables}")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # ---------- Component 8: Page numbers starting after TOC (0.10 points) ----------
    # Golden has 3 sections, PAGE field in footer, and pgNumType start=1 on section 2
    try:
        num_sections = len(doc.sections)
        has_page_field = False
        has_page_restart = False

        for s in doc.sections:
            # Check footer for PAGE field
            if s.footer and s.footer.paragraphs:
                for fp in s.footer.paragraphs:
                    xml = fp._element.xml
                    if 'PAGE' in xml and 'fldChar' in xml:
                        has_page_field = True

            # Check for page number restart (pgNumType with start attribute)
            sectPr = s._sectPr
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                pg_start = pgNumType.get(qn('w:start'))
                if pg_start is not None:
                    has_page_restart = True

        if num_sections >= 2 and has_page_field and has_page_restart:
            print(f"PASS: Component 8 — {num_sections} sections, PAGE field, page restart found (0.10 pts)")
            total_score += 0.10
        elif num_sections >= 2 and has_page_field:
            print(f"PARTIAL: Component 8 — {num_sections} sections, PAGE field but no restart (0.05 pts)")
            total_score += 0.05
        elif has_page_field:
            print(f"PARTIAL: Component 8 — PAGE field found but single section (0.03 pts)")
            total_score += 0.03
        else:
            print(f"FAIL: Component 8 — sections={num_sections}, page_field={has_page_field}, restart={has_page_restart}")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    # ---------- Component 9: Title page formatting (0.05 points) ----------
    # Golden has centered, bold, large title on first page
    try:
        if len(doc.paragraphs) < 3:
            print(f"FAIL: Component 9 — Too few paragraphs for title page")
        else:
            title_centered = False
            title_bold = False
            title_large = False

            # Check first non-empty paragraph (company name or title)
            for p in doc.paragraphs[:8]:
                if not p.text.strip():
                    continue
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    title_centered = True
                for r in p.runs:
                    if r.font.bold:
                        title_bold = True
                    if r.font.size and r.font.size > Pt(16):
                        title_large = True
                # Only need to find at least one centered+bold+large para in the title area
                if title_centered and title_bold and title_large:
                    break

            if title_centered and title_bold and title_large:
                print(f"PASS: Component 9 — Title page: centered, bold, large font (0.05 pts)")
                total_score += 0.05
            elif title_centered and (title_bold or title_large):
                print(f"PARTIAL: Component 9 — Title page partial: centered={title_centered}, bold={title_bold}, large={title_large} (0.02 pts)")
                total_score += 0.02
            else:
                print(f"FAIL: Component 9 — Title page: centered={title_centered}, bold={title_bold}, large={title_large}")
    except Exception as e:
        print(f"ERROR: Component 9 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
