"""
Initial Setup: Insert a date field at the top of a meeting agenda document
Task ID: writer_struct_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_struct_008'
OUTPUT = f'{WORKDIR}/team_meeting.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Heading 1: Meeting Agenda ---
    heading = doc.add_heading('Meeting Agenda', level=1)

    # --- 5 bullet-point agenda items ---
    agenda_items = [
        'Q2 Budget Review and Financial Projections',
        'Product Roadmap Updates and Milestone Discussion',
        'Team Performance Metrics and KPI Assessment',
        'Client Partnership Opportunities and Onboarding Status',
        'Action Items, Deadlines, and Next Steps',
    ]

    for item in agenda_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
