"""
Initial Setup: Enable Word Completion with min word length 5 in LibreOffice Writer
Task ID: writer_edit_044
Domain: libreoffice_writer

This script:
1. Creates a 20-page technical manual docx at ~/Desktop/long_document.docx
2. Ensures LibreOffice Word Completion is DISABLED (default state)
3. Opens the document in LibreOffice Writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_044'
OUTPUT = f'{WORKDIR}/Desktop/long_document.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def ensure_desktop_dir():
    """Ensure the Desktop directory exists."""
    desktop = f'{WORKDIR}/Desktop'
    os.makedirs(desktop, exist_ok=True)


def reset_word_completion_settings():
    """
    Ensure Word Completion is DISABLED in LibreOffice settings.
    LibreOffice stores AutoCorrect/Word Completion settings in registrymodifications.xcu.
    We remove any existing Word Completion enable entry so it reverts to the default (disabled).
    """
    config_path = f'{WORKDIR}/.config/libreoffice/4/user/registrymodifications.xcu'

    if not os.path.exists(config_path):
        # Config doesn't exist yet — defaults apply (Word Completion disabled by default)
        return

    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Remove any lines that enable Word Completion
        # The key is: /org.openoffice.Office.Writer/AutoFunction/Text/WordCompletion with
        # item name="IsActive" value="true" or MinWordLen entries
        import re

        # Remove the entire item node for WordCompletion IsActive if set to true
        # Pattern: <item oor:path="...WordCompletion"><prop oor:name="IsActive"...>true<...>
        content_new = re.sub(
            r'<item oor:path="/org\.openoffice\.Office\.Writer/AutoFunction/Text/WordCompletion">\s*'
            r'<prop oor:name="IsActive"[^/]*/>\s*</item>\n?',
            '',
            content
        )

        # Also handle multi-line
        content_new = re.sub(
            r'<item oor:path="/org\.openoffice\.Office\.Writer/AutoFunction/Text/WordCompletion">\s*'
            r'<prop oor:name="IsActive".*?</prop>\s*</item>\n?',
            '',
            content_new,
            flags=re.DOTALL
        )

        # Remove MinWordLen entries (reset to default)
        content_new = re.sub(
            r'<item oor:path="/org\.openoffice\.Office\.Writer/AutoFunction/Text/WordCompletion">\s*'
            r'<prop oor:name="MinWordLen".*?</prop>\s*</item>\n?',
            '',
            content_new,
            flags=re.DOTALL
        )

        with open(config_path, 'w', encoding='utf-8') as f:
            f.write(content_new)

        print(f'Reset Word Completion settings in {config_path}')

    except Exception as e:
        print(f'Warning: Could not reset Word Completion settings: {e}')


def create_initial():
    ensure_desktop_dir()

    doc = Document()

    # Title page
    title = doc.add_heading('Industrial Automation Systems: Technical Reference Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle_para = doc.add_paragraph('Version 3.2 | Revision Date: March 2025')
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    dept_para = doc.add_paragraph('Engineering Documentation Division\nApplied Technologies Corporation')
    dept_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction to Industrial Automation ............. 3',
        '2. System Architecture Overview ...................... 4',
        '3. Hardware Components and Specifications ............ 5',
        '4. Software Integration Framework ................... 7',
        '5. Network Communication Protocols .................. 8',
        '6. Safety and Compliance Standards .................. 10',
        '7. Installation and Commissioning ................... 12',
        '8. Calibration Procedures ........................... 14',
        '9. Maintenance and Troubleshooting .................. 16',
        '10. Appendices and Reference Data ................... 18',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # Chapter 1
    doc.add_heading('Chapter 1: Introduction to Industrial Automation', level=1)
    doc.add_paragraph(
        'Industrial automation represents the use of control systems—such as computers, robots, and '
        'information technologies—to handle different processes and machinery in an industry to replace '
        'a human being. It is the second step beyond mechanization in the scope of industrialization. '
        'Earlier, the purpose of automation was to increase productivity (since automated systems can '
        'work 24 hours a day), and to reduce the cost associated with human operators.'
    )
    doc.add_paragraph(
        'Modern industrial automation systems integrate programmable logic controllers (PLCs), '
        'distributed control systems (DCS), supervisory control and data acquisition (SCADA) systems, '
        'and human-machine interfaces (HMIs) into cohesive operational frameworks. These systems '
        'enable real-time monitoring, precise process control, and data-driven decision-making across '
        'manufacturing facilities worldwide.'
    )

    doc.add_heading('1.1 Historical Development', level=2)
    doc.add_paragraph(
        'The history of industrial automation dates back to the early 20th century when the first '
        'feedback control mechanisms were developed. James Watt\'s centrifugal governor (1788) is '
        'often cited as one of the earliest examples of automatic control. The development of relay '
        'logic in the 1940s and the invention of the transistor in 1947 marked pivotal moments that '
        'would eventually lead to the sophisticated automation systems we rely on today.'
    )
    doc.add_paragraph(
        'The introduction of the Modicon 084 PLC in 1968 by Dick Morley revolutionized manufacturing '
        'automation. This device replaced the complex relay wiring with a programmable solid-state '
        'device that could be easily reprogrammed without rewiring. The automotive industry was an '
        'early adopter, and the technology quickly spread to chemical processing, oil refining, and '
        'power generation sectors.'
    )

    doc.add_heading('1.2 Current State of the Industry', level=2)
    doc.add_paragraph(
        'Today, industrial automation encompasses a broad spectrum of technologies. The emergence of '
        'Industry 4.0—the fourth industrial revolution—has brought about the integration of cyber-'
        'physical systems, the Internet of Things (IoT), and cloud computing into manufacturing '
        'environments. Smart factories leverage machine learning algorithms and big data analytics '
        'to optimize production processes in ways previously impossible.'
    )

    doc.add_page_break()

    # Chapter 2
    doc.add_heading('Chapter 2: System Architecture Overview', level=1)
    doc.add_paragraph(
        'A modern industrial automation system is typically structured in hierarchical layers, each '
        'serving distinct functions while communicating with adjacent layers through standardized '
        'interfaces. Understanding this architecture is fundamental to effective system design, '
        'integration, and maintenance.'
    )

    doc.add_heading('2.1 Purdue Reference Model', level=2)
    doc.add_paragraph(
        'The Purdue Enterprise Reference Architecture (PERA), developed by Theodore J. Williams in '
        'the 1990s, provides a widely-adopted framework for understanding industrial control system '
        'hierarchy. This model defines five primary levels:'
    )

    levels_table = doc.add_table(rows=6, cols=2)
    levels_table.style = 'Table Grid'
    levels_table.cell(0, 0).text = 'Level'
    levels_table.cell(0, 1).text = 'Description'
    levels_table.cell(1, 0).text = 'Level 0: Field Devices'
    levels_table.cell(1, 1).text = 'Sensors, actuators, and basic I/O devices directly interfacing with physical processes'
    levels_table.cell(2, 0).text = 'Level 1: Basic Control'
    levels_table.cell(2, 1).text = 'PLCs, DCS controllers managing individual process units'
    levels_table.cell(3, 0).text = 'Level 2: Area Supervisory'
    levels_table.cell(3, 1).text = 'SCADA systems, HMIs providing operator interface and area supervision'
    levels_table.cell(4, 0).text = 'Level 3: Site Manufacturing'
    levels_table.cell(4, 1).text = 'MES, production scheduling, quality management systems'
    levels_table.cell(5, 0).text = 'Level 4: Enterprise'
    levels_table.cell(5, 1).text = 'ERP, business intelligence, financial and supply chain management'

    doc.add_page_break()

    # Chapter 3
    doc.add_heading('Chapter 3: Hardware Components and Specifications', level=1)
    doc.add_paragraph(
        'This chapter provides detailed specifications for the primary hardware components used in '
        'the AT-3200 Series automation platform. All components are designed to operate in '
        'industrial environments meeting IEC 61010-1 safety standards and IP65 enclosure ratings '
        'for dust and water resistance.'
    )

    doc.add_heading('3.1 Central Processing Unit Module', level=2)
    doc.add_paragraph(
        'The AT-3200 CPU Module (Part No. CPU-3200-EX) serves as the primary controller for the '
        'automation platform. Built on a dual-core ARM Cortex-A9 processor running at 1.2 GHz, '
        'it provides deterministic real-time control with scan times as low as 0.5 milliseconds.'
    )

    # Specs table
    specs = [
        ('Processor', 'Dual-core ARM Cortex-A9 @ 1.2 GHz'),
        ('RAM', '512 MB DDR3 ECC'),
        ('Flash Storage', '4 GB eMMC'),
        ('Communication', 'Dual Gigabit Ethernet, USB 2.0 x2'),
        ('I/O Expansion', 'Up to 32 expansion modules'),
        ('Operating Temp.', '-25°C to +70°C'),
        ('Power Supply', '24V DC ±20%'),
        ('Current Draw', '450 mA maximum'),
        ('Dimensions', '50mm × 130mm × 110mm (W×H×D)'),
        ('Certifications', 'CE, UL 508, cULus, ATEX Zone 2'),
    ]

    spec_table = doc.add_table(rows=len(specs) + 1, cols=2)
    spec_table.style = 'Table Grid'
    spec_table.cell(0, 0).text = 'Parameter'
    spec_table.cell(0, 1).text = 'Value'
    for i, (param, val) in enumerate(specs, 1):
        spec_table.cell(i, 0).text = param
        spec_table.cell(i, 1).text = val

    doc.add_page_break()

    # Chapter 4
    doc.add_heading('Chapter 4: Software Integration Framework', level=1)
    doc.add_paragraph(
        'The AT-3200 platform uses an open software architecture based on IEC 61131-3 programming '
        'standards. This enables engineers familiar with industry-standard PLC programming languages '
        'to develop and deploy applications without extensive retraining. The framework supports all '
        'five IEC 61131-3 languages: Ladder Diagram (LD), Structured Text (ST), Function Block '
        'Diagram (FBD), Sequential Function Chart (SFC), and Instruction List (IL).'
    )

    doc.add_heading('4.1 Development Environment', level=2)
    doc.add_paragraph(
        'AT Studio 4.0 provides the integrated development environment for the AT-3200 platform. '
        'Key features include syntax highlighting, auto-completion, online debugging with variable '
        'monitoring, cross-reference tools, and integrated version control via Git. The IDE runs '
        'on Windows 10/11 and Ubuntu 20.04 LTS, requiring a minimum of 8 GB RAM and 20 GB disk space.'
    )

    doc.add_page_break()

    # Chapter 5
    doc.add_heading('Chapter 5: Network Communication Protocols', level=1)
    doc.add_paragraph(
        'Industrial network communication requires protocols that prioritize determinism, reliability, '
        'and real-time performance over the flexibility and throughput prioritized in enterprise IT '
        'networks. The AT-3200 platform supports multiple industry-standard protocols to ensure '
        'compatibility with existing plant infrastructure.'
    )

    protocols = [
        ('EtherNet/IP', 'Standard', 'IEEE 802.3', '100/1000 Mbps', 'Cycle time ≥ 1 ms'),
        ('PROFINET IRT', 'Standard', 'IEEE 802.3', '100 Mbps', 'Cycle time ≥ 250 μs'),
        ('Modbus TCP', 'Standard', 'TCP/IP', 'Up to 100 Mbps', 'Cycle time ≥ 5 ms'),
        ('OPC UA', 'Standard', 'TCP/IP', 'Up to 1 Gbps', 'Cycle time ≥ 100 ms'),
        ('PROFIBUS DP', 'Fieldbus', 'RS-485', '12 Mbps', 'Cycle time ≥ 1 ms'),
        ('CANopen', 'Fieldbus', 'CAN bus', '1 Mbps', 'Cycle time ≥ 1 ms'),
    ]

    proto_table = doc.add_table(rows=len(protocols) + 1, cols=5)
    proto_table.style = 'Table Grid'
    headers_row = ['Protocol', 'Type', 'Physical', 'Max Speed', 'Min Cycle']
    for j, h in enumerate(headers_row):
        proto_table.cell(0, j).text = h
    for i, row_data in enumerate(protocols, 1):
        for j, val in enumerate(row_data):
            proto_table.cell(i, j).text = val

    doc.add_page_break()

    # Chapter 6
    doc.add_heading('Chapter 6: Safety and Compliance Standards', level=1)
    doc.add_paragraph(
        'Safety is paramount in industrial automation. The AT-3200 platform is designed to meet '
        'and exceed international safety standards including IEC 61508 (Functional Safety of '
        'Electrical/Electronic/Programmable Electronic Safety-related Systems) and ISO 13849 '
        '(Safety of Machinery - Safety-related Parts of Control Systems).'
    )

    doc.add_heading('6.1 Safety Integrity Levels', level=2)
    doc.add_paragraph(
        'The Safety Integrity Level (SIL) defines the reliability requirement for a safety '
        'instrumented function. The AT-3200 Safety CPU module achieves SIL 3 certification, '
        'meaning it can be used in applications where the probability of dangerous failure per '
        'hour is between 10⁻⁸ and 10⁻⁷. This qualifies it for use in critical applications '
        'including emergency shutdown systems, fire and gas detection, and burner management.'
    )

    doc.add_page_break()

    # Chapter 7
    doc.add_heading('Chapter 7: Installation and Commissioning', level=1)
    doc.add_paragraph(
        'Proper installation and commissioning of the AT-3200 system is critical to achieving '
        'optimal performance and ensuring long-term reliability. This chapter provides step-by-step '
        'procedures for mechanical installation, electrical wiring, software configuration, and '
        'initial system testing.'
    )

    doc.add_heading('7.1 Pre-Installation Checklist', level=2)
    checklist_items = [
        'Verify all equipment matches the bill of materials (BOM)',
        'Inspect all components for shipping damage',
        'Confirm mounting location meets environmental specifications',
        'Ensure adequate clearance for ventilation (minimum 50mm on all sides)',
        'Verify power supply voltage and current capacity',
        'Confirm network infrastructure is in place',
        'Review and obtain all required permits',
        'Assemble required tools: DIN rail cutters, wire strippers, torque screwdriver',
    ]
    for item in checklist_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Chapter 8
    doc.add_heading('Chapter 8: Calibration Procedures', level=1)
    doc.add_paragraph(
        'Accurate calibration is essential for maintaining measurement precision and process '
        'reliability. The AT-3200 platform includes built-in calibration utilities accessible '
        'through AT Studio or the on-device web interface. Calibration records are stored with '
        'timestamps, technician ID, and calibration certificate numbers for traceability.'
    )

    doc.add_heading('8.1 Analog Input Calibration', level=2)
    doc.add_paragraph(
        'Analog input modules (AI-3200-8V, AI-3200-8I) require calibration using reference '
        'standards traceable to NIST or equivalent national metrology institute. Calibration '
        'intervals depend on the application criticality: SIL applications require annual '
        'calibration; standard applications may extend to 24-month intervals with demonstrated '
        'measurement stability.'
    )

    doc.add_paragraph(
        'Calibration procedure for 4-20 mA current inputs:\n'
        '1. Connect a precision current source (accuracy ±0.01% FS) to the channel under test\n'
        '2. Apply 4.000 mA and record the displayed value\n'
        '3. Apply 20.000 mA and record the displayed value\n'
        '4. If either reading deviates by more than 0.1% FS, initiate two-point calibration\n'
        '5. Follow the on-screen wizard to apply offset and gain corrections\n'
        '6. Verify calibration with a third reference point at 12.000 mA (midscale)\n'
        '7. Record all measurements in the calibration log'
    )

    doc.add_page_break()

    # Chapter 9
    doc.add_heading('Chapter 9: Maintenance and Troubleshooting', level=1)
    doc.add_paragraph(
        'Preventive maintenance is the foundation of high-availability industrial automation. '
        'The AT-3200 platform provides comprehensive diagnostic tools including built-in '
        'oscilloscopes for signal analysis, network traffic monitoring, and predictive '
        'maintenance algorithms that analyze operational trends to forecast component failures '
        'before they occur.'
    )

    doc.add_heading('9.1 Scheduled Maintenance Tasks', level=2)
    maintenance_items = [
        ('Monthly', 'Inspect cabinet ventilation filters; clean or replace if blocked'),
        ('Monthly', 'Check LED status indicators on all modules; investigate any anomalies'),
        ('Quarterly', 'Verify all terminal connections for tightness (torque to specification)'),
        ('Quarterly', 'Review system event log for recurring warnings or errors'),
        ('Semi-annual', 'Test UPS battery backup; replace if capacity below 80%'),
        ('Annual', 'Full I/O functional test with process disconnected'),
        ('Annual', 'Calibration verification for all analog channels'),
        ('Annual', 'Firmware update review and application if required'),
        ('3-Year', 'Replace battery-backed RAM module battery'),
        ('5-Year', 'Fan replacement in CPU module and power supply'),
    ]

    maint_table = doc.add_table(rows=len(maintenance_items) + 1, cols=2)
    maint_table.style = 'Table Grid'
    maint_table.cell(0, 0).text = 'Interval'
    maint_table.cell(0, 1).text = 'Task'
    for i, (interval, task) in enumerate(maintenance_items, 1):
        maint_table.cell(i, 0).text = interval
        maint_table.cell(i, 1).text = task

    doc.add_page_break()

    # Chapter 10
    doc.add_heading('Chapter 10: Appendices and Reference Data', level=1)

    doc.add_heading('Appendix A: Error Codes', level=2)
    doc.add_paragraph(
        'The AT-3200 system uses a structured error code system. Error codes are displayed on '
        'the CPU module front panel and reported via OPC UA, SNMP traps, and email notifications. '
        'All error events are logged with timestamp, severity level, and context information.'
    )

    error_codes = [
        ('E001', 'Critical', 'CPU overtemperature', 'Check ventilation, verify ambient temperature'),
        ('E002', 'Critical', 'Power supply failure', 'Check input voltage, replace power supply module'),
        ('E003', 'Warning', 'Battery low', 'Replace battery module within 30 days'),
        ('E004', 'Warning', 'Communication timeout', 'Check network cables and switch configuration'),
        ('E005', 'Info', 'Configuration change', 'Log only; review change if unintended'),
        ('E006', 'Critical', 'I/O module fault', 'Identify module by LED; replace or re-seat module'),
        ('E007', 'Warning', 'Analog input out of range', 'Check sensor and wiring; verify scaling'),
        ('E008', 'Warning', 'Watchdog timeout', 'Review program scan time; check for infinite loops'),
        ('E009', 'Info', 'User login', 'Security audit log entry'),
        ('E010', 'Critical', 'Firmware checksum error', 'Restore firmware from known-good backup'),
    ]

    error_table = doc.add_table(rows=len(error_codes) + 1, cols=4)
    error_table.style = 'Table Grid'
    error_headers = ['Code', 'Severity', 'Description', 'Recommended Action']
    for j, h in enumerate(error_headers):
        error_table.cell(0, j).text = h
    for i, row_data in enumerate(error_codes, 1):
        for j, val in enumerate(row_data):
            error_table.cell(i, j).text = val

    doc.add_page_break()

    doc.add_heading('Appendix B: Spare Parts List', level=2)
    doc.add_paragraph(
        'The following spare parts are recommended for maintaining high system availability. '
        'Quantities are based on a typical installation of one CPU module, eight I/O modules, '
        'and two power supplies. Adjust quantities based on your specific system configuration '
        'and criticality requirements.'
    )

    spare_parts = [
        ('CPU-3200-EX', 'CPU Module', '1', 'Cold spare for critical applications'),
        ('PS-3200-24V', '24VDC Power Supply', '1', 'Hot-swap capable; keep on-site'),
        ('DI-3200-16', 'Digital Input Module', '2', 'Most common I/O type'),
        ('DO-3200-16R', 'Digital Output Relay Module', '2', 'High-wear due to relay cycling'),
        ('AI-3200-8V', 'Analog Input Voltage', '1', 'For calibration during repair'),
        ('BAT-3200', 'SRAM Battery', '4', '3-year replacement cycle'),
        ('FAN-3200', 'CPU Fan Assembly', '2', '5-year replacement cycle'),
        ('CF-16GB', 'CompactFlash Card', '2', 'Program and data backup'),
        ('FUSE-10A', '10A Blade Fuse', '10', 'Consumable; keep multiple'),
        ('CABLE-ETH-3M', 'Cat6 Patch Cable 3m', '5', 'For service connections'),
    ]

    spare_table = doc.add_table(rows=len(spare_parts) + 1, cols=4)
    spare_table.style = 'Table Grid'
    spare_headers = ['Part Number', 'Description', 'Qty', 'Notes']
    for j, h in enumerate(spare_headers):
        spare_table.cell(0, j).text = h
    for i, row_data in enumerate(spare_parts, 1):
        for j, val in enumerate(row_data):
            spare_table.cell(i, j).text = val

    doc.add_page_break()

    doc.add_heading('Appendix C: Glossary of Terms', level=2)
    glossary = [
        ('CIP', 'Common Industrial Protocol — communication standard underlying EtherNet/IP and DeviceNet'),
        ('DCS', 'Distributed Control System — control system architecture using distributed processing'),
        ('FMEA', 'Failure Mode and Effects Analysis — systematic analysis of potential failure modes'),
        ('HMI', 'Human Machine Interface — operator interface panel or software for process visualization'),
        ('IEC', 'International Electrotechnical Commission — standards organization for electrical technologies'),
        ('MES', 'Manufacturing Execution System — production management software layer'),
        ('OPC UA', 'Open Platform Communications Unified Architecture — platform-independent data exchange standard'),
        ('PFD', 'Probability of Failure on Demand — key SIL metric for safety systems'),
        ('PLC', 'Programmable Logic Controller — industrial digital computer for automation control'),
        ('SCADA', 'Supervisory Control and Data Acquisition — system for monitoring distributed assets'),
        ('SIL', 'Safety Integrity Level — IEC 61508 discrete level for specifying safety system requirements'),
        ('UPS', 'Uninterruptible Power Supply — backup power system preventing equipment shutdown on power loss'),
    ]

    glossary_table = doc.add_table(rows=len(glossary) + 1, cols=2)
    glossary_table.style = 'Table Grid'
    glossary_table.cell(0, 0).text = 'Abbreviation'
    glossary_table.cell(0, 1).text = 'Definition'
    for i, (abbr, defn) in enumerate(glossary, 1):
        glossary_table.cell(i, 0).text = abbr
        glossary_table.cell(i, 1).text = defn

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Reset Word Completion to disabled (default state)
    reset_word_completion_settings()

    # GUI-ready startup: open the document in LibreOffice Writer
    # Kill any existing LibreOffice instances first to ensure clean start
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(1.0)

    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
