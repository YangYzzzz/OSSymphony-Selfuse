"""
FINAL REWARD SCRIPT - SUCCESS
Task: Make the introduction single-spaced, set the body to double spacing, and align the conclusion left at 1.5 spacing, using 12 pt across.
Generated: 2025-10-14 06:06:14
Status: success
Model: azure-o3
Total Steps: 1
"""

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
import os, math

def spacing_matches(paragraph, target):
    """Check if paragraph line-spacing matches target (single, double, 1.5)."""
    fmt = paragraph.paragraph_format
    rule = fmt.line_spacing_rule
    spacing = fmt.line_spacing  # may be Length / float / None

    # Match by explicit rule first
    if rule is not None:
        if target == 'single' and rule == WD_LINE_SPACING.SINGLE:
            return True
        if target == 'double' and rule == WD_LINE_SPACING.DOUBLE:
            return True
        if target == '1.5' and rule == WD_LINE_SPACING.ONE_POINT_FIVE:
            return True

    # Fallback: numeric value (1 = single, 2 = double, 1.5 = one-and-half)
    if spacing is not None:
        try:
            val = float(spacing)
            if target == 'single' and math.isclose(val, 1.0, rel_tol=0.1):
                return True
            if target == 'double' and math.isclose(val, 2.0, rel_tol=0.1):
                return True
            if target == '1.5' and math.isclose(val, 1.5, rel_tol=0.1):
                return True
        except Exception:
            pass
    return False

def verify_font_size(doc):
    """Ensure every run is 12 pt (or inherits it)."""
    invalid = 0
    total = 0
    for para in doc.paragraphs:
        for run in para.runs:
            total += 1
            size = run.font.size  # Length or None
            if size is None:
                continue  # inherits – assume document default (12 pt)
            if not math.isclose(size.pt, 12.0, abs_tol=0.5):
                invalid += 1
    return invalid == 0, total, invalid

def verify_task(file_path):
    print(f"Verifying document: {file_path}")
    score = 0.0

    # ----- prerequisite: file must load -----
    if not os.path.exists(file_path):
        print("✗ File not found")
        return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Error loading DOCX: {e}")
        return 0.0

    # ----- requirement 1: 12 pt throughout (0.4) -----
    font_ok, total_runs, invalid_runs = verify_font_size(doc)
    print(f"Font size check → total runs: {total_runs}, invalid: {invalid_runs}")
    if font_ok and total_runs > 0:
        print("✓ All text is 12 pt (0.4)")
        score += 0.4
    else:
        print("✗ Non-12 pt text detected (0 points)")

    # Prepare flags for remaining formatting checks
    intro_ok = False
    body_ok = False
    concl_ok = False

    # Iterate paragraphs once to evaluate spacing/alignment conditions
    for para in doc.paragraphs:
        txt = para.text.strip().lower()
        if not txt:
            continue

        # Introduction paragraph (keyword contains 'introduction') must be single-spaced
        if 'introduction' in txt and spacing_matches(para, 'single'):
            intro_ok = True

        # Conclusion paragraph requires left alignment & 1.5 spacing
        if 'conclusion' in txt and spacing_matches(para, '1.5') and (para.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)):
            concl_ok = True

        # Body paragraphs – anything that is *not* intro or conclusion – need double spacing
        if 'introduction' not in txt and 'conclusion' not in txt and spacing_matches(para, 'double'):
            body_ok = True

    # ----- requirement 2: introduction spacing (0.2) -----
    if intro_ok:
        print("✓ Introduction is single-spaced (0.2)")
        score += 0.2
    else:
        print("✗ Introduction single spacing missing (0 points)")

    # ----- requirement 3: body double spacing (0.2) -----
    if body_ok:
        print("✓ Body paragraphs are double-spaced (0.2)")
        score += 0.2
    else:
        print("✗ Body double spacing missing (0 points)")

    # ----- requirement 4: conclusion alignment & 1.5 spacing (0.2) -----
    if concl_ok:
        print("✓ Conclusion left-aligned with 1.5 spacing (0.2)")
        score += 0.2
    else:
        print("✗ Conclusion formatting incorrect (0 points)")

    final_score = min(score, 1.0)
    print(f"REWARD: {final_score}")
    return final_score

# ---------- Execute verification when script is run ----------
if __name__ == "__main__":
    FILE_PATH = "/home/user/make_the_introduction_single_spaced_set_the_body_to_double_spacing_and_align_the_conclusion_left_at_.docx"
    verify_task(FILE_PATH)

