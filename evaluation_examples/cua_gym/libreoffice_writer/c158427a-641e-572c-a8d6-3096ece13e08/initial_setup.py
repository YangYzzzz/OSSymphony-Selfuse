"""
Initial Setup: Safety manual document with 5 paragraphs, WARNING in paragraph 3 unformatted
Task ID: writer_txtfmt_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_txtfmt_019'
OUTPUT = f'{WORKDIR}/safety_manual.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Workplace Safety Guidelines')
    title_run.bold = True
    title_run.font.size = Pt(18)

    # Paragraph 1
    p1 = doc.add_paragraph()
    r1 = p1.add_run(
        'This document outlines the safety procedures and guidelines that all employees '
        'must follow while on company premises. Adherence to these guidelines is mandatory '
        'and essential for maintaining a safe working environment for everyone.'
    )
    r1.font.name = 'Arial'
    r1.font.size = Pt(12)

    # Paragraph 2
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        'All employees are required to complete the mandatory safety training within their '
        'first two weeks of employment. Refresher courses are scheduled annually. Employees '
        'must present their safety certification card upon entering restricted areas.'
    )
    r2.font.name = 'Arial'
    r2.font.size = Pt(12)

    # Paragraph 3 — WARNING word must NOT be bold and NOT red in initial state
    p3 = doc.add_paragraph()
    # 'WARNING' run — regular, black, 12pt Arial
    r3_warning = p3.add_run('WARNING')
    r3_warning.bold = False
    r3_warning.font.name = 'Arial'
    r3_warning.font.size = Pt(12)
    r3_warning.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # Rest of sentence
    r3_rest = p3.add_run(
        ': All personnel must wear protective equipment when entering Zone C.'
    )
    r3_rest.font.name = 'Arial'
    r3_rest.font.size = Pt(12)

    # Paragraph 4
    p4 = doc.add_paragraph()
    r4 = p4.add_run(
        'Emergency exits are clearly marked throughout all floors of the facility. '
        'Employees must familiarize themselves with the nearest emergency exit to their '
        'workstation. Monthly evacuation drills are conducted to ensure readiness.'
    )
    r4.font.name = 'Arial'
    r4.font.size = Pt(12)

    # Paragraph 5
    p5 = doc.add_paragraph()
    r5 = p5.add_run(
        'Any safety concerns or incidents must be reported immediately to the designated '
        'Safety Officer on duty. Failure to report incidents may result in disciplinary '
        'action. The safety hotline is available 24/7 at extension 5200.'
    )
    r5.font.name = 'Arial'
    r5.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
