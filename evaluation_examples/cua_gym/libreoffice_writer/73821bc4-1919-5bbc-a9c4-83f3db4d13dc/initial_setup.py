"""
Initial Setup: Create a thesis document with a Table of Contents heading
Task ID: writer_acad_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Title Page --
    for _ in range(4):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run("Machine Learning Approaches to Natural Language Understanding")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("by\nElena M. Rodriguez")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run("Department of Computer Science\nStanford University\nMarch 2025")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # -- Page break before Table of Contents --
    doc.add_page_break()

    # -- Table of Contents heading --
    # Use default "Contents Heading" style if available, otherwise plain paragraph
    # We intentionally do NOT use Heading 1 style and do NOT center it
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # NOT centered
    run = toc_heading.add_run("Table of Contents")  # NOT uppercase
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # -- Simulated TOC entries --
    toc_entries = [
        ("Chapter 1: Introduction", "1"),
        ("   1.1 Background and Motivation", "3"),
        ("   1.2 Research Questions", "5"),
        ("   1.3 Thesis Outline", "7"),
        ("Chapter 2: Literature Review", "9"),
        ("   2.1 Transformer Architectures", "10"),
        ("   2.2 Pre-training Strategies", "15"),
        ("   2.3 Fine-tuning Methods", "20"),
        ("Chapter 3: Methodology", "25"),
        ("   3.1 Dataset Collection", "26"),
        ("   3.2 Model Architecture", "30"),
        ("   3.3 Training Procedure", "35"),
        ("Chapter 4: Experiments and Results", "40"),
        ("   4.1 Experimental Setup", "41"),
        ("   4.2 Baseline Comparisons", "45"),
        ("   4.3 Ablation Studies", "50"),
        ("Chapter 5: Discussion", "55"),
        ("   5.1 Key Findings", "56"),
        ("   5.2 Limitations", "60"),
        ("Chapter 6: Conclusion", "63"),
        ("References", "67"),
        ("Appendices", "75"),
    ]

    for entry, page_num in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(f"{entry}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        tab_run = p.add_run(f"\t{page_num}")
        tab_run.font.size = Pt(12)
        tab_run.font.name = "Times New Roman"

    # -- Page break before Chapter 1 --
    doc.add_page_break()

    # -- Chapter 1 --
    ch1 = doc.add_heading("Chapter 1: Introduction", level=1)

    intro_text = (
        "Natural language understanding (NLU) remains one of the most challenging "
        "problems in artificial intelligence. Despite remarkable progress in recent years, "
        "current systems still struggle with nuanced semantic interpretation, contextual "
        "reasoning, and the inherent ambiguity of human language. This thesis investigates "
        "novel machine learning approaches that aim to bridge the gap between statistical "
        "pattern matching and genuine language comprehension."
    )
    p = doc.add_paragraph(intro_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_heading("1.1 Background and Motivation", level=2)

    bg_text = (
        "The emergence of transformer-based architectures, particularly BERT and GPT, "
        "has fundamentally reshaped the landscape of NLU research. These models, pre-trained "
        "on vast corpora of text data, have demonstrated an unprecedented ability to capture "
        "linguistic patterns across multiple levels of abstraction. However, their reliance "
        "on surface-level statistical regularities raises important questions about the "
        "nature and depth of the understanding they achieve."
    )
    p = doc.add_paragraph(bg_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    motivation_text = (
        "Our work is motivated by three key observations: (1) existing models often fail "
        "on adversarial examples that require genuine reasoning, (2) performance gains from "
        "scaling alone appear to be plateauing for certain task categories, and (3) there is "
        "a growing need for interpretable NLU systems in high-stakes applications such as "
        "legal analysis, medical diagnosis, and scientific discovery."
    )
    p = doc.add_paragraph(motivation_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_heading("1.2 Research Questions", level=2)

    rqs = [
        "How can hierarchical attention mechanisms improve compositional understanding in transformer models?",
        "What role does explicit syntactic structure play in enhancing semantic representation quality?",
        "Can curriculum learning strategies improve sample efficiency for domain-specific NLU tasks?",
    ]
    for i, rq in enumerate(rqs, 1):
        p = doc.add_paragraph(f"RQ{i}: {rq}")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # -- Chapter 2 (partial) --
    doc.add_page_break()
    doc.add_heading("Chapter 2: Literature Review", level=1)

    lit_text = (
        "This chapter provides a comprehensive review of the relevant literature spanning "
        "three main areas: transformer architectures, pre-training strategies, and "
        "fine-tuning methods for downstream NLU tasks."
    )
    p = doc.add_paragraph(lit_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_heading("2.1 Transformer Architectures", level=2)

    transformer_text = (
        "The transformer architecture, introduced by Vaswani et al. (2017), relies entirely "
        "on self-attention mechanisms to process sequential data. Unlike recurrent neural "
        "networks, transformers can process all positions in a sequence simultaneously, "
        "enabling significantly more efficient training on modern GPU hardware. The "
        "multi-head attention mechanism allows the model to jointly attend to information "
        "from different representation subspaces at different positions."
    )
    p = doc.add_paragraph(transformer_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
