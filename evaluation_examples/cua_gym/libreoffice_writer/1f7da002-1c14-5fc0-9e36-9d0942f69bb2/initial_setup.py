"""
Initial Setup: Create a Writer document with a 4-column, 6-row vendor comparison table.
Task ID: writer_biz_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title paragraph above the table
    title = doc.add_heading('Q3 2025 Procurement Report', level=1)

    intro = doc.add_paragraph(
        'The following table summarizes our evaluation of four shortlisted vendors '
        'for the enterprise cloud migration project. Scores reflect weighted criteria '
        'across technical capability, cost efficiency, support responsiveness, and '
        'contract flexibility.'
    )

    # Create a 4-column, 6-row table (row 0 = headers, rows 1-5 = data)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # Row 0: individual header cells (NOT merged - task is to merge them)
    headers = ['Vendor Name', 'Technical Score', 'Cost Rating', 'Overall Rank']
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Data rows with realistic vendor comparison data
    data = [
        ['Meridian Cloud Solutions',  '87.5',  'A-',   '2'],
        ['NovaTech Infrastructure',   '91.2',  'B+',   '1'],
        ['Pinnacle Systems Group',    '78.8',  'A',    '3'],
        ['Stratos Digital Services',  '74.3',  'A+',   '4'],
        ['Vertex Global Partners',    '69.1',  'B',    '5'],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value

    # Add a note paragraph after the table
    doc.add_paragraph('')
    note = doc.add_paragraph(
        'Note: Technical scores are computed from the weighted rubric distributed on '
        '2025-07-14. Cost ratings follow the A+ through D scale established by the '
        'Finance Committee. Final vendor selection is pending board approval scheduled '
        'for September 8, 2025.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
