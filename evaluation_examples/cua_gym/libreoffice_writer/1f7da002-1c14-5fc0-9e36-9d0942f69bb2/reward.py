"""
Reward Script: Merge first row cells into spanning header 'Vendor Comparison Analysis'
Task ID: writer_biz_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): First row cells are merged (gridSpan=4)
  Component 2 (0.3): Merged cell text is 'Vendor Comparison Analysis'
  Component 3 (0.2): Text is center-aligned
  Component 4 (0.2): Text is bold and 14pt
"""

import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_050'


def persist_app_state(domain):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for %s" % domain)
    except Exception as e:
        print("PERSIST_WARN: save hook failed: %s" % e)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table has at least 1 row
    if len(table.rows) == 0:
        print("CRITICAL: Table has no rows")
        print("REWARD: 0.0")
        return 0.0

    row0 = table.rows[0]

    # Component 1: First row cells are merged (gridSpan=4) (0.3 points)
    try:
        first_cell_tc = row0.cells[0]._tc
        tcPr = first_cell_tc.find(qn('w:tcPr'))
        merged = False
        if tcPr is not None:
            gs = tcPr.find(qn('w:gridSpan'))
            if gs is not None:
                val = gs.get(qn('w:val'))
                if val is not None and int(val) >= 4:
                    merged = True

        # Also verify all cells in row 0 point to the same tc element
        all_same = all(row0.cells[i]._tc is row0.cells[0]._tc for i in range(1, len(row0.cells)))

        if merged and all_same:
            print("PASS: Component 1 -- First row cells merged with gridSpan=%s (0.3 pts)" % val)
            total_score += 0.3
        elif merged:
            print("PARTIAL: Component 1 -- gridSpan found but cells not identical objects (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 1 -- First row cells not merged (gridSpan not found or < 4)")
    except Exception as e:
        print("ERROR: Component 1 -- %s" % e)

    # Component 2: Merged cell text is 'Vendor Comparison Analysis' (0.3 points)
    try:
        cell_text = row0.cells[0].text.strip()
        if cell_text == 'Vendor Comparison Analysis':
            print("PASS: Component 2 -- Cell text is 'Vendor Comparison Analysis' (0.3 pts)")
            total_score += 0.3
        elif 'vendor comparison' in cell_text.lower():
            print("PARTIAL: Component 2 -- Text contains 'vendor comparison' but not exact: '%s' (0.15 pts)" % cell_text)
            total_score += 0.15
        else:
            print("FAIL: Component 2 -- Expected 'Vendor Comparison Analysis', found: '%s'" % cell_text)
    except Exception as e:
        print("ERROR: Component 2 -- %s" % e)

    # Component 3: Text is center-aligned (0.2 points)
    try:
        cell_paras = row0.cells[0].paragraphs
        # Check alignment of the paragraph containing the text
        centered = False
        for para in cell_paras:
            if para.text.strip():
                align = para.paragraph_format.alignment
                if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    centered = True
                    break
        if centered:
            print("PASS: Component 3 -- Text is center-aligned (0.2 pts)")
            total_score += 0.2
        else:
            print("FAIL: Component 3 -- Text is not center-aligned")
    except Exception as e:
        print("ERROR: Component 3 -- %s" % e)

    # Component 4: Text is bold and 14pt (0.2 points)
    try:
        bold_found = False
        size_14_found = False
        for para in row0.cells[0].paragraphs:
            for run in para.runs:
                if run.text.strip():
                    if run.font.bold is True:
                        bold_found = True
                    if run.font.size is not None and abs(run.font.size.pt - 14.0) < 0.5:
                        size_14_found = True

        if bold_found and size_14_found:
            print("PASS: Component 4 -- Text is bold and 14pt (0.2 pts)")
            total_score += 0.2
        else:
            # No partial credit: bold alone exists in initial state (precondition),
            # and 14pt alone without bold doesn't meet the task requirement.
            details = "bold=%s, size_14=%s" % (bold_found, size_14_found)
            print("FAIL: Component 4 -- Need both bold AND 14pt (%s)" % details)
    except Exception as e:
        print("ERROR: Component 4 -- %s" % e)

    final_score = min(total_score, 1.0)
    print("\nScore: %.1f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Persist any unsaved GUI state, then verify
persist_app_state('libreoffice_writer')

file_path = os.path.join(WORKDIR, TASK_ID + '.docx')
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
