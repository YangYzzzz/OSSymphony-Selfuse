"""
Initial Setup: Court order document with no page borders
Task ID: writer_legal_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_072'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Court Header ---
    h = doc.add_paragraph()
    h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run("IN THE SUPERIOR COURT OF THE STATE OF CALIFORNIA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    h2 = doc.add_paragraph()
    h2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h2.paragraph_format.space_after = Pt(12)
    run2 = h2.add_run("FOR THE COUNTY OF LOS ANGELES")
    run2.bold = True
    run2.font.size = Pt(13)
    run2.font.name = "Times New Roman"

    # --- Case Number ---
    cn = doc.add_paragraph()
    cn.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cn.paragraph_format.space_after = Pt(18)
    run_cn = cn.add_run("Case No. 2025-CV-04871")
    run_cn.font.size = Pt(12)
    run_cn.font.name = "Times New Roman"
    run_cn.bold = True

    # --- Parties ---
    parties = doc.add_paragraph()
    parties.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    parties.paragraph_format.space_after = Pt(6)
    r = parties.add_run("ELENA VASQUEZ,")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    p_line = doc.add_paragraph()
    p_line.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_line.paragraph_format.space_after = Pt(2)
    r2 = p_line.add_run("Plaintiff,")
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"
    r2.italic = True

    vs = doc.add_paragraph()
    vs.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    vs.paragraph_format.space_after = Pt(6)
    r3 = vs.add_run("v.")
    r3.font.size = Pt(12)
    r3.font.name = "Times New Roman"

    def_name = doc.add_paragraph()
    def_name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    def_name.paragraph_format.space_after = Pt(2)
    r4 = def_name.add_run("REDWOOD CAPITAL PARTNERS, LLC,")
    r4.font.size = Pt(12)
    r4.font.name = "Times New Roman"

    def_label = doc.add_paragraph()
    def_label.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    def_label.paragraph_format.space_after = Pt(18)
    r5 = def_label.add_run("Defendant.")
    r5.font.size = Pt(12)
    r5.font.name = "Times New Roman"
    r5.italic = True

    # --- Title ---
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(18)
    rt = title.add_run("ORDER GRANTING PLAINTIFF'S MOTION FOR SUMMARY JUDGMENT")
    rt.bold = True
    rt.font.size = Pt(13)
    rt.font.name = "Times New Roman"
    rt.underline = True

    # --- Body Paragraphs ---
    def add_body(text, space_after=Pt(10)):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = space_after
        p.paragraph_format.first_line_indent = Inches(0.5)
        r = p.add_run(text)
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"
        return p

    add_body(
        "THIS MATTER came before the Court on Plaintiff Elena Vasquez's Motion for "
        "Summary Judgment filed on January 15, 2025. The Court, having reviewed the "
        "motion papers, the opposition filed by Defendant Redwood Capital Partners, LLC "
        "on February 3, 2025, and the reply thereto, and having heard oral arguments on "
        "February 28, 2025, and being fully advised in the premises, hereby enters the "
        "following Order:"
    )

    # --- Findings ---
    findings_title = doc.add_paragraph()
    findings_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    findings_title.paragraph_format.space_before = Pt(12)
    findings_title.paragraph_format.space_after = Pt(10)
    rf = findings_title.add_run("FINDINGS OF FACT")
    rf.bold = True
    rf.font.size = Pt(12)
    rf.font.name = "Times New Roman"
    rf.underline = True

    findings = [
        "1. On or about March 12, 2023, Plaintiff entered into a Consulting Services "
        "Agreement (the \"Agreement\") with Defendant for financial advisory services "
        "related to the acquisition of Meridian Technologies, Inc.",

        "2. Under Section 4.2 of the Agreement, Defendant agreed to provide monthly "
        "performance reports and maintain fiduciary standards consistent with industry "
        "practices as outlined in the California Business and Professions Code § 17200.",

        "3. Between April 2023 and September 2024, Defendant failed to deliver seven (7) "
        "of the eighteen (18) required monthly reports, constituting a material breach of "
        "the Agreement.",

        "4. As a direct and proximate result of Defendant's breach, Plaintiff sustained "
        "financial damages in the amount of $342,750.00, which includes lost investment "
        "returns of $278,500.00 and additional consulting fees of $64,250.00 paid to "
        "replacement advisors.",

        "5. Defendant has not presented any genuine issue of material fact to dispute the "
        "breach or the calculation of damages.",
    ]

    for f in findings:
        add_body(f)

    # --- Conclusions of Law ---
    conclusions_title = doc.add_paragraph()
    conclusions_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    conclusions_title.paragraph_format.space_before = Pt(12)
    conclusions_title.paragraph_format.space_after = Pt(10)
    rc = conclusions_title.add_run("CONCLUSIONS OF LAW")
    rc.bold = True
    rc.font.size = Pt(12)
    rc.font.name = "Times New Roman"
    rc.underline = True

    conclusions = [
        "1. This Court has jurisdiction over this matter pursuant to California Code of "
        "Civil Procedure § 437c.",

        "2. There are no triable issues of material fact as required under Aguilar v. "
        "Atlantic Richfield Co. (2001) 25 Cal.4th 826.",

        "3. Plaintiff is entitled to judgment as a matter of law on her breach of "
        "contract claim.",
    ]

    for c in conclusions:
        add_body(c)

    # --- Order ---
    order_title = doc.add_paragraph()
    order_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    order_title.paragraph_format.space_before = Pt(12)
    order_title.paragraph_format.space_after = Pt(10)
    ro = order_title.add_run("ORDER")
    ro.bold = True
    ro.font.size = Pt(12)
    ro.font.name = "Times New Roman"
    ro.underline = True

    add_body(
        "IT IS HEREBY ORDERED, ADJUDGED, AND DECREED that Plaintiff's Motion for "
        "Summary Judgment is GRANTED. Defendant Redwood Capital Partners, LLC shall pay "
        "to Plaintiff Elena Vasquez the sum of Three Hundred Forty-Two Thousand Seven "
        "Hundred Fifty Dollars ($342,750.00), together with pre-judgment interest at the "
        "rate of 10% per annum from September 30, 2024, to the date of this Order, and "
        "post-judgment interest at the legal rate from the date of this Order until paid "
        "in full."
    )

    add_body(
        "IT IS FURTHER ORDERED that Defendant shall bear all costs of this action, "
        "including reasonable attorney's fees as determined by subsequent motion."
    )

    # --- Signature Block ---
    doc.add_paragraph()  # spacing

    sig = doc.add_paragraph()
    sig.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sig.paragraph_format.space_after = Pt(2)
    rs = sig.add_run("Dated: March 14, 2025")
    rs.font.size = Pt(12)
    rs.font.name = "Times New Roman"

    doc.add_paragraph()  # spacing for signature

    judge_line = doc.add_paragraph()
    judge_line.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    judge_line.paragraph_format.space_after = Pt(2)
    rj = judge_line.add_run("_________________________________")
    rj.font.size = Pt(12)
    rj.font.name = "Times New Roman"

    judge_name = doc.add_paragraph()
    judge_name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    judge_name.paragraph_format.space_after = Pt(2)
    rjn = judge_name.add_run("Hon. Patricia M. Nakamura")
    rjn.font.size = Pt(12)
    rjn.font.name = "Times New Roman"
    rjn.bold = True

    judge_title_p = doc.add_paragraph()
    judge_title_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    rjt = judge_title_p.add_run("Judge of the Superior Court")
    rjt.font.size = Pt(12)
    rjt.font.name = "Times New Roman"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
