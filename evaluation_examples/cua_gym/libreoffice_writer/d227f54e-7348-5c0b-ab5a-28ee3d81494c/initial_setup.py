"""
Initial Setup: Insert a 4x6 table into quarterly sales report
Task ID: writer_tm_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Heading
    doc.add_heading('Quarterly Sales Summary', level=1)

    # Realistic paragraph about Q3 sales performance
    para_text = (
        "The third quarter of 2025 demonstrated strong momentum across all major product "
        "categories. Total revenue reached $2.87 million, representing a 14.3% increase "
        "compared to Q2 figures. Our enterprise solutions division led the growth with a "
        "23% quarter-over-quarter improvement, driven primarily by new client acquisitions "
        "in the healthcare and financial services sectors. Consumer product lines maintained "
        "steady performance, with subscription renewals exceeding 91% for the second "
        "consecutive quarter. The Southeast Asian market expansion contributed an additional "
        "$340,000 in revenue, surpassing initial projections by 18%. Looking ahead, the "
        "sales team has identified several high-value opportunities in the pipeline that "
        "are expected to close in Q4, positioning the company well for year-end targets."
    )
    doc.add_paragraph(para_text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
