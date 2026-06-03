"""
Reward Script: Insert a 4x6 table into a quarterly sales report document
Task ID: writer_tm_001
Domain: libreoffice_writer
Scoring:
  Component 1 (0.40) — Table exists in the document
  Component 2 (0.30) — Table has exactly 4 columns and 6 rows
  Component 3 (0.15) — All table cells are empty
  Component 4 (0.15) — Table uses 'Table Grid' style (default borders)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_001'


def persist_app_state(domain: str):
    """Best-effort save of any unsaved LibreOffice edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: heading and paragraph still exist (gate, not scored)
    if len(doc.paragraphs) < 2:
        print("FAIL: Precondition — document should have at least 2 paragraphs (heading + body)")
        print("REWARD: 0.0")
        return 0.0

    heading_ok = doc.paragraphs[0].text.strip().startswith("Quarterly Sales Summary")
    if not heading_ok:
        print(f"FAIL: Precondition — expected heading 'Quarterly Sales Summary', found: {doc.paragraphs[0].text[:60]!r}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: A table exists in the document (0.40 points)
    try:
        num_tables = len(doc.tables)
        if num_tables >= 1:
            print(f"PASS: Component 1 — Document contains {num_tables} table(s) (0.40 pts)")
            total_score += 0.40
        else:
            print(f"FAIL: Component 1 — No tables found in the document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # If no table, remaining components cannot be checked
    if len(doc.tables) == 0:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    table = doc.tables[0]

    # Component 2: Table has exactly 4 columns and 6 rows (0.30 points)
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows == 6 and num_cols == 4:
            print(f"PASS: Component 2 — Table dimensions are 6 rows x 4 cols (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — Expected 6 rows x 4 cols, found {num_rows} rows x {num_cols} cols")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All table cells are empty (0.15 points)
    try:
        non_empty_cells = []
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if cell.text.strip():
                    non_empty_cells.append((ri, ci, cell.text.strip()))

        if len(non_empty_cells) == 0:
            print(f"PASS: Component 3 — All table cells are empty (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — {len(non_empty_cells)} non-empty cell(s): {non_empty_cells[:5]}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Table uses 'Table Grid' style (default borders) (0.15 points)
    try:
        style_name = table.style.name if table.style else None
        if style_name == "Table Grid":
            print(f"PASS: Component 4 — Table style is 'Table Grid' (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 — Expected style 'Table Grid', found: {style_name!r}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
