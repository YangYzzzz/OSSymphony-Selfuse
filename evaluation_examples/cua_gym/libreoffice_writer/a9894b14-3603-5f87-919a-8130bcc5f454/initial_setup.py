"""
Initial Setup: Sales Award document with data source references but no conditional fields.
Task ID: writer_mt_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('Quarterly Sales Recognition Award', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Pinnacle Solutions Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.bold = True

    doc.add_paragraph()  # spacing

    # --- Date line ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = date_para.add_run('Date: March 31, 2025')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # --- Salutation ---
    salutation = doc.add_paragraph()
    run = salutation.add_run('Dear Sales Representative,')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # --- Body paragraph 1 ---
    body1 = doc.add_paragraph()
    run = body1.add_run(
        'We are pleased to announce the results of our quarterly sales performance review. '
        'Each representative has been evaluated based on their assigned region and total sales amount '
        'for the period. The following recognition category has been determined for your performance:'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # --- Placeholder for where the conditional field should go ---
    award_para = doc.add_paragraph()
    run = award_para.add_run('Award Category: ')
    run.bold = True
    run.font.size = Pt(12)
    run2 = award_para.add_run('[INSERT CONDITIONAL FIELD HERE]')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_paragraph()  # spacing

    # --- Body paragraph 2 ---
    body2 = doc.add_paragraph()
    run = body2.add_run(
        'This recognition is based on the data from our SalesData records, which include your '
        'SalesRepName, Region, SalesAmount, and Quarter information. The classification criteria are as follows:'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # --- Criteria list ---
    criteria = [
        'If your Region is "North" and your SalesAmount exceeds $10,000, you are classified as "Top Performer - North".',
        'If your Region is "North" and your SalesAmount is $10,000 or less, you are classified as "North Region".',
        'All other regions are classified as "Other Region".',
    ]
    for c in criteria:
        p = doc.add_paragraph(c, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    # --- Data source reference table ---
    heading2 = doc.add_heading('Data Source Reference: SalesData', level=2)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['SalesRepName', 'Region', 'SalesAmount', 'Quarter']
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['Elena Marchetti', 'North', '$47,200', 'Q1 2025'],
        ['James Okafor', 'South', '$31,850', 'Q1 2025'],
        ['Priya Sharma', 'North', '$8,400', 'Q1 2025'],
        ['Carlos Mendez', 'East', '$22,100', 'Q1 2025'],
        ['Aisha Williams', 'West', '$15,750', 'Q1 2025'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    # --- Closing ---
    closing1 = doc.add_paragraph()
    run = closing1.add_run(
        'We appreciate your dedication and hard work throughout this quarter. '
        'Please direct any questions regarding your classification to the Sales Operations team.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    closing2 = doc.add_paragraph()
    run = closing2.add_run('Best regards,')
    run.font.size = Pt(11)

    sig = doc.add_paragraph()
    run = sig.add_run('Victoria Langston')
    run.bold = True
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    run = sig_title.add_run('Director of Sales Operations')
    run.font.size = Pt(11)

    sig_company = doc.add_paragraph()
    run = sig_company.add_run('Pinnacle Solutions Inc.')
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
