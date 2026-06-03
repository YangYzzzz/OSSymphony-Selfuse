"""
Reward Script: Direct Mail Piece for Top-Tier Customers
Task ID: writer_mktg_055
Domain: libreoffice_writer
Scoring:
  - Component 1: Header "Apex Dynamics" bold 18pt                        (0.20 pts)
  - Component 2: Salutation contains {Customer_Name} merge field          (0.15 pts)
  - Component 3: Offer box in table with yellow bg and dark blue border   (0.25 pts)
  - Component 4: Dashed line separator present (centered)                 (0.15 pts)
  - Component 5: "RESPONSE CARD" heading bold 14pt                        (0.15 pts)
  - Component 6: Response fields with underlined blanks                   (0.10 pts)
  Total: 1.0
"""

import os
import re
import lxml.etree as etree

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_055'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Header contains "Apex Dynamics" as bold 18pt text (0.20 pts)
    # The initial document has an empty header; golden adds 'Apex Dynamics' bold 18pt.
    # -------------------------------------------------------------------------
    try:
        section = doc.sections[0]
        header_paras = section.header.paragraphs
        apex_in_header = any('Apex Dynamics' in p.text for p in header_paras)
        # Count runs where 'Apex Dynamics' appears with correct 18pt size
        runs_with_correct_size = sum(
            1 for para in header_paras
            for run in para.runs
            if 'Apex Dynamics' in run.text
            and run.font.size is not None
            and abs(run.font.size.pt - 18.0) < 1.0
        )
        if apex_in_header and runs_with_correct_size >= 1:
            print(f"PASS: Component 1 — Header 'Apex Dynamics' found with 18pt formatting (0.20 pts)")
            total_score += 0.20
        elif apex_in_header and runs_with_correct_size == 0:
            print(f"PASS (partial): Component 1 — Header 'Apex Dynamics' present but size not 18pt (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — 'Apex Dynamics' not found in document header")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Salutation paragraph contains {Customer_Name} placeholder (0.15 pts)
    # The initial document has 'Dear Valued Customer,'; golden has 'Dear {Customer_Name},'.
    # -------------------------------------------------------------------------
    try:
        customer_name_count = sum(
            1 for p in doc.paragraphs
            if '{Customer_Name}' in p.text and 'Dear' in p.text
        )
        if customer_name_count >= 1:
            print(f"PASS: Component 2 — Salutation with '{{Customer_Name}}' found (0.15 pts)")
            total_score += 0.15
        else:
            # Relax: check presence without 'Dear'
            any_placeholder = sum(1 for p in doc.paragraphs if '{Customer_Name}' in p.text)
            if any_placeholder >= 1:
                print(f"PASS (partial): Component 2 — '{{Customer_Name}}' found but salutation format differs (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 2 — '{{Customer_Name}}' placeholder not found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Offer box in a table with yellow background (#FFF9C4) and
    # dark blue border (#003366), containing offer text bold 16pt (0.25 pts)
    # -------------------------------------------------------------------------
    try:
        offer_text_target = '50% off annual Enterprise plan'
        offer_cells_found = [
            cell
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
            if offer_text_target.lower() in cell.text.lower()
        ]

        if not offer_cells_found:
            print(f"FAIL: Component 3 — Offer text not found in any table cell")
        else:
            cell = offer_cells_found[0]
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr_xml = etree.tostring(tcPr, encoding=str).upper()

            has_yellow_bg = 'FFF9C4' in tcPr_xml
            has_blue_border = '003366' in tcPr_xml

            if has_yellow_bg and has_blue_border:
                print(f"PASS: Component 3 — Offer box in table, yellow bg (#FFF9C4), dark blue border (#003366) (0.25 pts)")
                total_score += 0.25
            elif has_yellow_bg or has_blue_border:
                print(f"PASS (partial): Component 3 — Offer in table, yellow={has_yellow_bg} blue_border={has_blue_border} (0.15 pts)")
                total_score += 0.15
            elif not has_yellow_bg and not has_blue_border:
                print(f"PASS (minimal): Component 3 — Offer in table but no matching bg/border colors (0.10 pts)")
                total_score += 0.10
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Dashed line separator present (0.15 pts)
    # Initial has no dashes; golden adds a centered "- - - - -" line.
    # -------------------------------------------------------------------------
    try:
        dashed_paras = [
            p for p in doc.paragraphs
            if re.search(r'([-–—]\s*){5,}', p.text.strip())
            or '- - - -' in p.text
        ]
        if dashed_paras:
            dash_text = dashed_paras[0].text.strip()[:40]
            print(f"PASS: Component 4 — Dashed separator line found: {dash_text!r} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 — Dashed separator line not found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: "RESPONSE CARD" heading bold 14pt (0.15 pts)
    # Initial has no RESPONSE CARD section; golden adds it after the dashed line.
    # -------------------------------------------------------------------------
    try:
        response_card_paras = [
            p for p in doc.paragraphs
            if 'RESPONSE CARD' in p.text.upper()
        ]
        if not response_card_paras:
            print(f"FAIL: Component 5 — 'RESPONSE CARD' heading not found")
        else:
            para = response_card_paras[0]
            # Check for bold 14pt in any run
            runs_with_response_card = [
                r for r in para.runs
                if 'RESPONSE' in r.text.upper() or 'CARD' in r.text.upper()
            ]
            formatting_correct = any(
                r.bold is True and r.font.size is not None and abs(r.font.size.pt - 14.0) < 1.0
                for r in runs_with_response_card
            )
            if formatting_correct:
                print(f"PASS: Component 5 — 'RESPONSE CARD' heading bold 14pt found (0.15 pts)")
                total_score += 0.15
            elif not formatting_correct:
                print(f"PASS (partial): Component 5 — 'RESPONSE CARD' present but formatting differs (0.08 pts)")
                total_score += 0.08
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -------------------------------------------------------------------------
    # Component 6: Response fields have underlined blanks (Name, Company, Email, Phone) (0.10 pts)
    # Initial has plain text fields; golden adds underlined blanks after each.
    # -------------------------------------------------------------------------
    try:
        fields = ['Name', 'Company', 'Email', 'Phone']
        fields_with_underline = sum(
            1 for field in fields
            for para in doc.paragraphs
            if para.text.strip().startswith(field + ':')
            and any(r.underline is True and r.text.strip() for r in para.runs)
        )
        if fields_with_underline >= 3:
            print(f"PASS: Component 6 — {fields_with_underline}/4 response fields have underlined blanks (0.10 pts)")
            total_score += 0.10
        elif fields_with_underline >= 1:
            pts = round(0.10 * fields_with_underline / 4, 4)
            print(f"PASS (partial): Component 6 — {fields_with_underline}/4 response fields with underlines ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 6 — No response fields with underlined blanks found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
