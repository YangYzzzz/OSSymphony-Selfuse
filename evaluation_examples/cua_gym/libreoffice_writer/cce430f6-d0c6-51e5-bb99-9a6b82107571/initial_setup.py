"""
Initial Setup: Direct mail piece for top-tier customers (plain text initial state)
Task ID: writer_mktg_055
Domain: libreoffice_writer

Creates the initial .docx with plain text only — no merge fields, no header,
no offer box, no dashed line, no formatted response section.
All text is 12pt, single column.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_plain_para(doc, text, font_size_pt=12, bold=False, alignment=None):
    """Add a paragraph with plain formatting (no special layout)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    if alignment is not None:
        para.paragraph_format.alignment = alignment
    return para


def create_initial():
    doc = Document()

    # Set up page margins (Letter size, 1-inch margins)
    section = doc.sections[0]
    from docx.shared import Inches
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Salutation (plain, no merge field) ---
    add_plain_para(doc, 'Dear Valued Customer,', font_size_pt=12)

    # --- Three offer description paragraphs ---
    add_plain_para(
        doc,
        (
            'We are thrilled to bring you an exclusive opportunity designed specifically '
            'for our most valued enterprise clients. As a preferred partner, you have helped '
            'shape the direction of our platform, and we want to give back in a meaningful way.'
        ),
        font_size_pt=12,
    )

    add_plain_para(
        doc,
        (
            'Our Enterprise plan delivers unlimited seats, priority 24/7 support, dedicated '
            'onboarding assistance, and advanced analytics dashboards — everything your team '
            'needs to scale confidently. Thousands of businesses already rely on our platform '
            'to streamline operations and drive measurable results.'
        ),
        font_size_pt=12,
    )

    add_plain_para(
        doc,
        (
            'For a limited time, we are extending a special renewal incentive exclusively to '
            'clients like you who have demonstrated outstanding engagement and loyalty. '
            'Do not miss this chance to lock in significant savings on the plan your team counts on.'
        ),
        font_size_pt=12,
    )

    # --- Offer details line (plain text, no box/frame/formatting) ---
    add_plain_para(
        doc,
        '50% off annual Enterprise plan \u2014 offer expires April 30, 2026',
        font_size_pt=12,
    )

    # --- Response fields (plain text, no underline formatting) ---
    for field in ['Name:', 'Company:', 'Email:', 'Phone:']:
        add_plain_para(doc, field, font_size_pt=12)

    # --- Reply instruction paragraph ---
    add_plain_para(
        doc,
        (
            'To claim your discount, simply complete and return this card by April 30, 2026. '
            'Mail it to Apex Dynamics, 100 Innovation Drive, San Francisco, CA 94105, '
            'or fax to 1-800-555-0199. Our team will contact you within two business days '
            'to confirm your renewed subscription at the discounted rate.'
        ),
        font_size_pt=12,
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
