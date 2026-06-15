"""
Reward Script: Set specific border styles on a table in custom_borders.docx
Task ID: writer_tbl_049
Domain: libreoffice_writer
Scoring:
  Component 1: Top border is 2pt (sz=16) solid black          — 0.25 pts
  Component 2: Bottom border is 2pt (sz=16) solid black       — 0.25 pts
  Component 3: Horizontal inner borders are 0.5pt dashed gray — 0.25 pts
  Component 4: Left, right, and vertical inner borders removed — 0.25 pts
  Total: 1.0

OOXML sz unit is 1/8th of a point, so:
  0.5pt => sz=4
  2pt   => sz=16
Gray in OOXML: 808080
"""

import os
from docx import Document
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_049'

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_border_attribs(tblBorders_elem, border_name):
    """
    Extract val, sz, and color attributes from a named border element.
    Returns a dict with keys 'val', 'sz', 'color' (all strings or None).
    """
    if tblBorders_elem is None:
        return None
    border_elem = tblBorders_elem.find(f'{{{NS}}}{border_name}')
    if border_elem is None:
        return None
    return {
        'val':   border_elem.get(f'{{{NS}}}val'),
        'sz':    border_elem.get(f'{{{NS}}}sz'),
        'color': border_elem.get(f'{{{NS}}}color'),
    }


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Checks that the table's tblBorders element reflects the required border
    changes:
      - top/bottom: 2pt solid black
      - insideH: 0.5pt dashed gray (#808080)
      - left/right/insideV: none (removed)
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: exactly one table with 4 rows and 3 columns
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Get tblBorders element from table properties
    tblPr = table._element.find(f'{{{NS}}}tblPr')
    if tblPr is None:
        print("FAIL: No tblPr element found — cannot verify borders")
        print("REWARD: 0.0")
        return 0.0

    tblBorders = tblPr.find(f'{{{NS}}}tblBorders')
    if tblBorders is None:
        print("FAIL: No tblBorders element found in tblPr")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Top border is 2pt (sz=16) solid black (0.25 points)
    try:
        top = get_border_attribs(tblBorders, 'top')
        if top is None:
            print("FAIL: Component 1 — top border element not found")
        else:
            # sz=16 means 2pt (OOXML: sz in eighths-of-a-point)
            val_ok    = top['val'] == 'single'
            sz_ok     = top['sz'] == '16'
            color_ok  = top['color'] is not None and top['color'].upper() == '000000'
            if val_ok and sz_ok and color_ok:
                print(f"PASS: Component 1 — top border is 2pt solid black (val={top['val']}, sz={top['sz']}, color={top['color']}) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — top border expected val=single/sz=16/color=000000, "
                      f"found val={top['val']}/sz={top['sz']}/color={top['color']}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Bottom border is 2pt (sz=16) solid black (0.25 points)
    try:
        bottom = get_border_attribs(tblBorders, 'bottom')
        if bottom is None:
            print("FAIL: Component 2 — bottom border element not found")
        else:
            val_ok    = bottom['val'] == 'single'
            sz_ok     = bottom['sz'] == '16'
            color_ok  = bottom['color'] is not None and bottom['color'].upper() == '000000'
            if val_ok and sz_ok and color_ok:
                print(f"PASS: Component 2 — bottom border is 2pt solid black (val={bottom['val']}, sz={bottom['sz']}, color={bottom['color']}) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — bottom border expected val=single/sz=16/color=000000, "
                      f"found val={bottom['val']}/sz={bottom['sz']}/color={bottom['color']}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Horizontal inner borders are 0.5pt dashed gray (0.25 points)
    # Gray = #808080; sz=4 means 0.5pt
    try:
        insideH = get_border_attribs(tblBorders, 'insideH')
        if insideH is None:
            print("FAIL: Component 3 — insideH border element not found")
        else:
            val_ok    = insideH['val'] == 'dashed'
            sz_ok     = insideH['sz'] == '4'
            color_ok  = insideH['color'] is not None and insideH['color'].upper() == '808080'
            if val_ok and sz_ok and color_ok:
                print(f"PASS: Component 3 — insideH is 0.5pt dashed gray (val={insideH['val']}, sz={insideH['sz']}, color={insideH['color']}) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — insideH expected val=dashed/sz=4/color=808080, "
                      f"found val={insideH['val']}/sz={insideH['sz']}/color={insideH['color']}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Left, right, and vertical inner borders are removed (val=none) (0.25 points)
    # All three must be none for full credit
    try:
        left     = get_border_attribs(tblBorders, 'left')
        right    = get_border_attribs(tblBorders, 'right')
        insideV  = get_border_attribs(tblBorders, 'insideV')

        left_none    = left    is not None and left['val']    == 'none'
        right_none   = right   is not None and right['val']   == 'none'
        insideV_none = insideV is not None and insideV['val'] == 'none'

        if left_none and right_none and insideV_none:
            print(f"PASS: Component 4 — left/right/insideV borders all removed (val=none) (0.25 pts)")
            total_score += 0.25
        else:
            details = []
            if not left_none:
                lv = left['val'] if left else 'missing'
                details.append(f"left={lv}")
            if not right_none:
                rv = right['val'] if right else 'missing'
                details.append(f"right={rv}")
            if not insideV_none:
                iv = insideV['val'] if insideV else 'missing'
                details.append(f"insideV={iv}")
            print(f"FAIL: Component 4 — expected all removed (none), but: {', '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
