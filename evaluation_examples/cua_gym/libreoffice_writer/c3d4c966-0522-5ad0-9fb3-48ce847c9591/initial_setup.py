"""
Initial Setup: Insert blank line after every sentence - Employee Policy Notice
Task ID: osworld_writer_blank_line_insertion_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_blank_line_insertion_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Document title
    title = doc.add_heading('Employee Policy Notice', level=1)

    # Paragraph 1: Attendance and Punctuality Policy (5 sentences)
    p1 = doc.add_paragraph(
        'All employees are expected to report to work on time and maintain consistent attendance throughout the year. '
        'Tardiness of more than fifteen minutes without prior notice will be recorded as an unexcused absence. '
        'Employees who anticipate being late must notify their direct supervisor at least thirty minutes before their scheduled start time. '
        'Repeated unexcused absences may result in disciplinary action, up to and including termination of employment. '
        'Any employee with an attendance concern should contact Human Resources to discuss available accommodations or leave options.'
    )

    # Paragraph 2: Workplace Conduct Policy (5 sentences)
    p2 = doc.add_paragraph(
        'All staff members are required to maintain a professional and respectful demeanor when interacting with colleagues, clients, and vendors. '
        'Harassment, discrimination, or any form of hostile behavior in the workplace is strictly prohibited and will not be tolerated. '
        'Employees who witness or experience inappropriate conduct should report it immediately to their manager or the Human Resources department. '
        'Confidential investigations will be conducted for all reported incidents to ensure a fair and impartial review process. '
        'Violations of the workplace conduct policy may result in immediate suspension or termination depending on the severity of the offense.'
    )

    # Paragraph 3: Remote Work and Equipment Policy (5 sentences)
    p3 = doc.add_paragraph(
        'Employees approved for remote work arrangements must maintain a secure and productive home office environment during all scheduled work hours. '
        'Company-issued equipment, including laptops, monitors, and peripherals, must be used exclusively for business-related activities. '
        'Any loss, damage, or theft of company equipment must be reported to the IT department within twenty-four hours of the incident. '
        'Remote employees are expected to participate in all scheduled virtual meetings and respond to communications within two business hours. '
        'The company reserves the right to revoke remote work privileges if performance standards or security requirements are not consistently met.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
