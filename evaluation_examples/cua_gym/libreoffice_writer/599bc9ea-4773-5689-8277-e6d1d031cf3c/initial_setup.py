"""
Initial Setup: Two-column newsletter with dense text, no hyphenation
Task ID: writer_rd_042
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_042'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: standard A4, narrow margins for newsletter ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # --- Set up two-column layout via section XML ---
    sectPr = section._sectPr
    cols_elem = sectPr.find(qn('w:cols'))
    if cols_elem is not None:
        sectPr.remove(cols_elem)
    cols = parse_xml(
        f'<w:cols {nsdecls("w")} w:num="2" w:space="360" w:equalWidth="1"/>'
    )
    sectPr.append(cols)

    # --- Explicitly suppress auto-hyphens on the document (ensure NO hyphenation) ---
    # We'll set suppressAutoHyphens on each paragraph to be safe

    # --- Newsletter Title ---
    title_para = doc.add_heading('The Metropolitan Chronicle', level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_para.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

    # Subtitle / date line
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = subtitle.add_run('Volume 14, Issue 3  |  March 2026  |  Community Edition')
    run_sub.font.size = Pt(9)
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run_sub.font.italic = True

    # Horizontal rule via bottom border on empty paragraph
    hr_para = doc.add_paragraph()
    hr_pPr = hr_para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6D"/>'
        f'</w:pBdr>'
    )
    hr_pPr.append(pBdr)

    # --- Dense paragraphs (justified, no hyphenation) ---
    paragraphs_content = [
        # Article 1
        ("Neighborhood Revitalization Project Gains Unprecedented Momentum",
         "heading"),
        ("The comprehensive neighborhood revitalization initiative, spearheaded by the "
         "Metropolitan Development Corporation in collaboration with community stakeholders, "
         "has demonstrated extraordinary progress throughout the southeastern districts of the "
         "metropolitan area. Commissioner Alexandra Whitfield-Patterson announced at yesterday's "
         "extraordinary council meeting that approximately $4.7 million in supplementary funding "
         "has been appropriated for infrastructure improvements, environmental remediation, and "
         "comprehensive transportation accessibility enhancements.", "body"),
        ("Representatives from the Neighborhood Environmental Sustainability Commission "
         "emphasized that the transformation encompasses substantially more than cosmetic "
         "improvements. The interconnected infrastructure modernization program includes "
         "underground utility replacement, environmentally sustainable stormwater management "
         "installations, photovoltaic-powered streetlighting, and universally accessible "
         "pedestrian thoroughfares incorporating bioretention landscaping elements.", "body"),
        ("Councilwoman Marguerite Fitzpatrick-Henderson, who has been an extraordinarily "
         "outspoken advocate for the disadvantaged neighborhoods, characterized the development "
         "as transformational. She underscored that approximately seventy-three percent of the "
         "appropriated funds would be directed toward historically underrepresented communities "
         "where infrastructure deterioration has been disproportionately concentrated.", "body"),

        # Article 2
        ("International Gastronomy Festival Celebrates Unprecedented Diversity",
         "heading"),
        ("The seventeenth annual International Gastronomy and Culinary Heritage Festival "
         "attracted an extraordinary congregation of approximately twenty-eight thousand "
         "enthusiastic participants over the three-day extravaganza held at the Waterfront "
         "Convention Center. Distinguished restaurateurs, internationally recognized culinary "
         "professionals, and accomplished home cooks representing forty-seven nationalities "
         "showcased their gastronomical traditions through demonstrations, collaborative "
         "workshops, and complimentary tastings.", "body"),
        ("Particularly noteworthy was the Mediterranean Fermentation Symposium, organized "
         "by microbiologist Dr. Konstantinos Papadimitriou, which investigated the "
         "biochemical transformations underlying traditional preservation methodologies "
         "including lacto-fermentation, acetification, and enzymatic decomposition. "
         "Attendees participated in hands-on experimentation with sourdough cultivation, "
         "traditional kimchi preparation, and artisanal cheesemaking incorporating "
         "thermophilic and mesophilic bacterial cultures.", "body"),

        # Article 3
        ("Community Preparedness Initiative Strengthens Emergency Response",
         "heading"),
        ("The Metropolitan Emergency Management Administration has inaugurated a "
         "comprehensive community preparedness certification program designed to "
         "strengthen neighborhood-level emergency response capabilities throughout "
         "the municipality. The multifaceted initiative encompasses standardized "
         "first aid and cardiopulmonary resuscitation training, neighborhood "
         "communication infrastructure development, and psychological preparedness "
         "workshops facilitated by experienced emergency management professionals.", "body"),
        ("Superintendent Christopher Worthington-MacAllister of the Metropolitan "
         "Fire and Rescue Department acknowledged the indispensable contribution of "
         "approximately three hundred and forty-seven community volunteers who have "
         "dedicated considerable personal time to the preliminary organizational "
         "phase. The department's comprehensive assessment determined that neighborhoods "
         "with established preparedness networks demonstrated substantially improved "
         "response coordination during the unprecedented thunderstorm emergencies "
         "experienced throughout the autumn months.", "body"),
        ("Furthermore, the administration has established interconnected communication "
         "hubs strategically distributed throughout twelve residential neighborhoods, "
         "incorporating battery-powered telecommunications equipment, comprehensive "
         "emergency supply repositories, and universally accessible information "
         "dissemination stations. These installations are specifically engineered to "
         "maintain operational functionality during extended electrical infrastructure "
         "interruptions and telecommunications network disruptions.", "body"),
    ]

    for text, ptype in paragraphs_content:
        if ptype == "heading":
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Georgia'
            # Set paragraph spacing
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)

        # Explicitly suppress auto-hyphens on every paragraph
        pPr = p._element.get_or_add_pPr()
        suppress = parse_xml(f'<w:suppressAutoHyphens {nsdecls("w")} w:val="1"/>')
        pPr.append(suppress)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
