"""
Initial Setup: Legal brief with body paragraphs needing first-line indent
Task ID: writer_legal_012
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_012'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # -- Default style: Normal (NO first-line indent) --
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 2.0
    # Explicitly ensure NO first-line indent
    style.paragraph_format.first_line_indent = Inches(0)

    # -- Title block (centered) --
    title_lines = [
        "IN THE UNITED STATES DISTRICT COURT",
        "FOR THE NORTHERN DISTRICT OF CALIFORNIA",
    ]
    for line in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # Blank separator
    doc.add_paragraph()

    # -- Case caption --
    caption_lines = [
        "HARRISON TECHNOLOGIES, INC.,",
        "    Plaintiff,",
        "",
        "        v.",
        "",
        "WESTERN PACIFIC SOLUTIONS, LLC,",
        "    Defendant.",
    ]
    for line in caption_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # Case number
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("Case No. 3:2025-cv-04817-JLR")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_paragraph()

    # -- Document title --
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("PLAINTIFF'S MEMORANDUM OF POINTS AND AUTHORITIES")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("IN SUPPORT OF MOTION FOR PRELIMINARY INJUNCTION")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_paragraph()

    # =================== SECTION I ===================
    h1 = doc.add_heading('I. INTRODUCTION', level=1)
    for run in h1.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    body_paragraphs = [
        "Plaintiff Harrison Technologies, Inc. (\"Harrison\") respectfully submits this "
        "memorandum in support of its motion for a preliminary injunction against Defendant "
        "Western Pacific Solutions, LLC (\"Western Pacific\"). The relief sought is necessary "
        "to prevent irreparable harm arising from Defendant's ongoing misappropriation of "
        "Plaintiff's proprietary trade secrets and confidential business information.",

        "As set forth below, Harrison has a strong likelihood of success on the merits of its "
        "claims under the Defend Trade Secrets Act, 18 U.S.C. \u00a7 1836, and the California "
        "Uniform Trade Secrets Act, Cal. Civ. Code \u00a7\u00a7 3426-3426.11. The balance of "
        "hardships tips sharply in Harrison's favor, and the public interest supports granting "
        "injunctive relief to protect lawful innovation and fair competition.",
    ]
    for text in body_paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # =================== SECTION II ===================
    h2 = doc.add_heading('II. STATEMENT OF FACTS', level=1)
    for run in h2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    facts = [
        "Harrison Technologies is a Delaware corporation headquartered in San Jose, California. "
        "Since its founding in 2014, Harrison has developed a suite of proprietary machine learning "
        "algorithms for predictive supply chain analytics, collectively known as the \"LogiSense\" "
        "platform. The development of LogiSense required an investment of over $23 million in "
        "research and development over an eight-year period.",

        "On or about March 15, 2024, Dr. Elena Vasquez, Harrison's former Director of Algorithm "
        "Development, resigned from her position at Harrison. Two weeks later, Dr. Vasquez accepted "
        "employment as Vice President of Data Science at Western Pacific Solutions, a direct "
        "competitor in the supply chain analytics market.",

        "Following Dr. Vasquez's departure, Harrison's internal security audit revealed that "
        "approximately 47,000 files containing source code, algorithm specifications, and customer "
        "data had been downloaded from Harrison's secure servers to an external storage device "
        "during the seventy-two hours preceding Dr. Vasquez's resignation. Digital forensic "
        "analysis confirmed that the downloads occurred under Dr. Vasquez's credentials.",

        "Within four months of Dr. Vasquez's departure, Western Pacific launched a competing "
        "product branded \"SupplyAI\" that exhibits substantial functional overlap with Harrison's "
        "LogiSense platform. Independent technical analysis conducted by Dr. Richard Tanaka of "
        "Stanford University concluded that SupplyAI employs algorithmic structures and optimization "
        "techniques that are \"strikingly similar\" to those used in LogiSense.",
    ]
    for text in facts:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # =================== SECTION III ===================
    h3 = doc.add_heading('III. LEGAL STANDARD', level=1)
    for run in h3.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    legal_paras = [
        "A party seeking a preliminary injunction must establish: (1) a likelihood of success "
        "on the merits; (2) a likelihood of irreparable harm in the absence of preliminary "
        "relief; (3) that the balance of equities tips in its favor; and (4) that an injunction "
        "is in the public interest. Winter v. Natural Resources Defense Council, Inc., 555 U.S. "
        "7, 20 (2008).",

        "The Ninth Circuit applies a \"sliding scale\" approach, under which a stronger showing "
        "on one element may offset a weaker showing on another, so long as there are \"serious "
        "questions going to the merits\" and the balance of hardships \"tips sharply\" toward the "
        "movant. Alliance for the Wild Rockies v. Cottrell, 632 F.3d 1127, 1135 (9th Cir. 2011).",
    ]
    for text in legal_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # =================== SECTION IV ===================
    h4 = doc.add_heading('IV. ARGUMENT', level=1)
    for run in h4.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # Sub-heading
    h4a = doc.add_heading('A. Harrison Is Likely to Succeed on the Merits', level=2)
    for run in h4a.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

    argument_paras = [
        "To establish a claim for trade secret misappropriation under the Defend Trade Secrets "
        "Act, a plaintiff must demonstrate that: (1) it possesses a trade secret; (2) the defendant "
        "misappropriated the trade secret; and (3) the trade secret was related to a product or "
        "service used in interstate commerce. 18 U.S.C. \u00a7 1836(b)(1).",

        "Harrison's LogiSense algorithms qualify as trade secrets under both federal and state "
        "law. Harrison has taken extensive measures to maintain the secrecy of its algorithms, "
        "including requiring all employees with access to sign non-disclosure agreements, "
        "implementing multi-factor authentication and access logging for its code repositories, "
        "and restricting access to LogiSense source code to a team of fewer than fifteen engineers.",

        "The evidence strongly supports a finding of misappropriation. The forensic evidence "
        "demonstrates that Dr. Vasquez downloaded tens of thousands of proprietary files "
        "immediately before departing Harrison. The subsequent launch of a substantially similar "
        "competing product by her new employer, within a timeline inconsistent with independent "
        "development, creates a compelling inference of misappropriation.",
    ]
    for text in argument_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    h4b = doc.add_heading('B. Harrison Will Suffer Irreparable Harm Without Injunctive Relief', level=2)
    for run in h4b.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

    harm_paras = [
        "Courts routinely recognize that the misappropriation of trade secrets gives rise to "
        "irreparable harm because the resulting injury is difficult to quantify and cannot be "
        "adequately compensated by monetary damages alone. Stuhlbarg Int'l Sales Co. v. John D. "
        "Brush & Co., 240 F.3d 832, 841 (9th Cir. 2001).",

        "Here, the continued use of Harrison's trade secrets by Western Pacific threatens to "
        "erode Harrison's competitive advantage in a market where first-mover advantages and "
        "proprietary technology are paramount. Each day that Western Pacific operates its SupplyAI "
        "platform using Harrison's misappropriated algorithms, Harrison loses market share that "
        "cannot be recovered through damages alone.",
    ]
    for text in harm_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # =================== SECTION V ===================
    h5 = doc.add_heading('V. CONCLUSION', level=1)
    for run in h5.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    conclusion = [
        "For the foregoing reasons, Harrison Technologies respectfully requests that this Court "
        "grant its motion for a preliminary injunction, enjoining Western Pacific Solutions from "
        "continuing to use, disclose, or benefit from Harrison's misappropriated trade secrets "
        "pending resolution of this action on the merits.",

        "A proposed order is submitted herewith.",
    ]
    for text in conclusion:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # -- Signature block --
    doc.add_paragraph()
    p = doc.add_paragraph("Respectfully submitted,")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("_________________________________")
    p = doc.add_paragraph()
    run = p.add_run("Katherine M. Thornton, Esq.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p = doc.add_paragraph("THORNTON & ASSOCIATES LLP")
    p = doc.add_paragraph("525 Market Street, Suite 3200")
    p = doc.add_paragraph("San Francisco, California 94105")
    p = doc.add_paragraph("Telephone: (415) 555-7890")
    p = doc.add_paragraph("Email: kthornton@thorntonlaw.com")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Attorneys for Plaintiff Harrison Technologies, Inc.")
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
