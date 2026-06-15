"""
Reward Script: Indent first line of body paragraphs by 0.5 inches in legal brief
Task ID: writer_legal_012
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Normal style first_line_indent == 457200 EMU (0.5 in / 1.27 cm)
  Component 2 (0.3): Body paragraphs effectively have ~0.5in first-line indent
  Component 3 (0.2): Heading paragraphs do NOT have first-line indent
"""

import os
from docx import Document
from docx.shared import Inches, Emu

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_012'
TARGET_INDENT = 457200  # 0.5 inches in EMU
TOLERANCE = 20000       # ~0.02 inch tolerance for rounding

def persist_app_state(domain):
    """Best-effort save via Ctrl+S in case document is open in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def get_effective_first_line_indent(para):
    """
    Return the effective first-line indent for a paragraph in EMU.
    If paragraph has explicit override, use that.
    Otherwise, inherit from style.
    """
    pf = para.paragraph_format
    if pf.first_line_indent is not None:
        return pf.first_line_indent
    # Inherit from style
    style = para.style
    if style and style.paragraph_format and style.paragraph_format.first_line_indent is not None:
        return style.paragraph_format.first_line_indent
    return 0


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Normal style first_line_indent is ~0.5 inches (0.5 points)
    # This is the PRIMARY change the task requires: modifying the Default Paragraph Style.
    # Initial state has Normal first_line_indent=0; golden has 457200 EMU (0.5in).
    try:
        normal_style = doc.styles['Normal']
        style_fli = normal_style.paragraph_format.first_line_indent
        if style_fli is not None and abs(style_fli - TARGET_INDENT) <= TOLERANCE:
            print(f"PASS: Component 1 — Normal style first_line_indent={style_fli} EMU (~0.5in) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Normal style first_line_indent={style_fli}, expected ~{TARGET_INDENT}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Body paragraphs effectively have ~0.5in first-line indent (0.3 points)
    # Identify body paragraphs: Normal-styled paragraphs that are actual brief content
    # (not caption, not signature block). We check that the majority of Normal-styled
    # paragraphs in the argument sections have effective indent close to 0.5in.
    # This catches cases where someone applied indent per-paragraph instead of via style.
    try:
        # Body paragraphs are those with Normal style that contain substantive text
        # and appear in the argument sections (after headings, before signature)
        body_paras = []
        in_body = False
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading'):
                in_body = True
                continue
            if in_body and style_name == 'Normal' and para.text.strip():
                # Stop collecting at signature block indicators
                text = para.text.strip()
                if text.startswith('Respectfully') or text.startswith('____'):
                    in_body = False
                    continue
                body_paras.append(para)

        if len(body_paras) == 0:
            print("FAIL: Component 2 — No body paragraphs found after headings")
        else:
            indented_count = 0
            for bp in body_paras:
                eff_indent = get_effective_first_line_indent(bp)
                if eff_indent is not None and abs(eff_indent - TARGET_INDENT) <= TOLERANCE:
                    indented_count += 1

            ratio = indented_count / len(body_paras)
            if ratio >= 0.8:
                print(f"PASS: Component 2 — {indented_count}/{len(body_paras)} body paragraphs have ~0.5in indent (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Only {indented_count}/{len(body_paras)} body paragraphs indented (need >=80%)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Heading paragraphs do NOT have 0.5in first-line indent (0.2 points)
    # Headings should remain unindented. This verifies that the indent was applied
    # specifically to body text, not blindly to everything.
    # In the initial state, headings also have no indent, so we need this check to be
    # specifically about headings NOT gaining indent when the style change is applied.
    # Since headings use Heading 1/2 styles (not Normal), they should be unaffected.
    # However, the task change (Normal style indent) only matters if headings are different.
    # We combine: headings must NOT have effective indent ~0.5in AND at least one body
    # paragraph must have it (to ensure this only passes on golden, not initial).
    try:
        heading_paras = [p for p in doc.paragraphs if p.style and p.style.name.startswith('Heading')]
        any_body_indented = any(
            abs(get_effective_first_line_indent(p) - TARGET_INDENT) <= TOLERANCE
            for p in body_paras
        ) if body_paras else False

        if not heading_paras:
            print("FAIL: Component 3 — No heading paragraphs found")
        else:
            headings_clean = True
            for hp in heading_paras:
                eff = get_effective_first_line_indent(hp)
                if eff is not None and abs(eff - TARGET_INDENT) <= TOLERANCE:
                    headings_clean = False
                    print(f"FAIL: Component 3 — Heading '{hp.text[:40]}' has ~0.5in indent (should not)")
                    break

            if headings_clean and any_body_indented:
                print(f"PASS: Component 3 — {len(heading_paras)} headings correctly unindented, body is indented (0.2 pts)")
                total_score += 0.2
            elif headings_clean and not any_body_indented:
                print(f"FAIL: Component 3 — Headings unindented but body also unindented (no task change)")
            else:
                print(f"FAIL: Component 3 — Some headings incorrectly indented")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'

persist_app_state("libreoffice_writer")

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
