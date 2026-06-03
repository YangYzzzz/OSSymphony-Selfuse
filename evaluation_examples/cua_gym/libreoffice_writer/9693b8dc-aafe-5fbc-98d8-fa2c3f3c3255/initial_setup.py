"""
Initial Setup: Insert a 'File Name' field in the footer
Task ID: writer_tm_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_069'
DOC_DIR = f'{WORKDIR}/Documents'
OUTPUT = f'{DOC_DIR}/Budget_2026.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DOC_DIR, exist_ok=True)

    doc = Document()

    # --- Title ---
    title = doc.add_heading("Annual Budget Report - Fiscal Year 2026", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report provides a comprehensive overview of the projected budget "
        "allocations for fiscal year 2026. The total organizational budget is "
        "estimated at $4.82 million, representing a 7.3% increase from the "
        "previous fiscal year. Key growth areas include technology infrastructure "
        "and talent acquisition."
    )

    # --- Department Budgets ---
    doc.add_heading("Department Budget Allocations", level=1)
    doc.add_paragraph(
        "The following table summarizes the approved budget for each department, "
        "along with the percentage change compared to FY2025."
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Department", "FY2025 Budget", "FY2026 Budget", "Change (%)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    data = [
        ["Engineering", "$1,250,000", "$1,410,000", "+12.8%"],
        ["Marketing", "$620,000", "$645,000", "+4.0%"],
        ["Sales", "$890,000", "$935,000", "+5.1%"],
        ["Human Resources", "$340,000", "$365,000", "+7.4%"],
        ["Finance & Accounting", "$285,000", "$298,000", "+4.6%"],
        ["Operations", "$510,000", "$548,000", "+7.5%"],
        ["Customer Support", "$410,000", "$432,000", "+5.4%"],
        ["Legal & Compliance", "$195,000", "$210,000", "+7.7%"],
    ]
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    # --- Capital Expenditures ---
    doc.add_heading("Capital Expenditures", level=1)
    doc.add_paragraph(
        "Major capital expenditure items for FY2026 include the following:"
    )
    items = [
        "Server infrastructure upgrade - $320,000 (Q1-Q2)",
        "Office renovation, Building C - $185,000 (Q2)",
        "Fleet vehicle replacement program - $94,000 (Q3)",
        "Security system modernization - $67,000 (Q1)",
        "Warehouse automation equipment - $142,000 (Q2-Q3)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # --- Revenue Projections ---
    doc.add_heading("Revenue Projections", level=1)
    doc.add_paragraph(
        "Based on current market analysis and sales pipeline data, projected "
        "revenue for FY2026 is $12.4 million, with a confidence interval of "
        "+/- 8%. The primary revenue drivers remain enterprise software licenses "
        "and professional services contracts. The consulting division is expected "
        "to contribute $2.1 million, up from $1.8 million in FY2025."
    )

    # --- Risk Assessment ---
    doc.add_heading("Risk Assessment", level=1)
    doc.add_paragraph(
        "Key financial risks identified for the upcoming fiscal year include "
        "potential supply chain disruptions in Q2, currency exchange volatility "
        "affecting international operations, and increased competition in the "
        "mid-market segment. Mitigation strategies have been developed for each "
        "scenario and are detailed in Appendix B."
    )

    # --- Approval Section ---
    doc.add_heading("Approval", level=1)
    doc.add_paragraph("Prepared by: Elena Vasquez, Senior Financial Analyst")
    doc.add_paragraph("Reviewed by: Thomas Nakamura, CFO")
    doc.add_paragraph("Date: March 15, 2026")

    # Footer is intentionally left EMPTY - the task is to add a filename field to it
    # Ensure footer exists but is empty
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear any default content
    for p in footer.paragraphs:
        p.text = ""

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
