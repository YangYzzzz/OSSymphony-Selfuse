"""
Reward Script: Verify professional offer letter for Michael Chen
Task ID: writer_hr_025
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Company name "Apex Digital Solutions" centered at top
  Component 2 (0.10): Date present
  Component 3 (0.15): Recipient address (Michael Chen, 456 Oak Avenue, Portland, OR 97201)
  Component 4 (0.10): Greeting (Dear Mr. Chen or Dear Michael)
  Component 5 (0.25): Offer details (Senior Software Engineer, $145,000, May 15, 2026)
  Component 6 (0.25): Closing with HR Director signature block
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_025'


def persist_app_state(domain):
    """Try to save any unsaved document via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def get_full_text(doc):
    """Get all paragraph text as a single lowercase string for searching."""
    return "\n".join(p.text for p in doc.paragraphs).lower()


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have content (at least a few paragraphs)
    if len(doc.paragraphs) < 3:
        print(f"FAIL: Document has only {len(doc.paragraphs)} paragraphs — appears blank or near-blank")
        print("REWARD: 0.0")
        return 0.0

    full_text = get_full_text(doc)

    # Component 1: Company name "Apex Digital Solutions" centered at top (0.15 points)
    try:
        centered_company_matches = [
            p for p in doc.paragraphs[:3]
            if "apex digital solutions" in p.text.lower()
            and p.paragraph_format.alignment in (WD_PARAGRAPH_ALIGNMENT.CENTER, 1)
        ]
        if len(centered_company_matches) > 0:
            print(f"PASS: Component 1 — Company name 'Apex Digital Solutions' found centered at top (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — 'Apex Digital Solutions' not found centered in first 3 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Date present (0.10 points)
    try:
        date_patterns = [
            r'\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}\b',
            r'\b\d{1,2}/\d{1,2}/\d{4}\b',
            r'\b\d{4}-\d{2}-\d{2}\b',
        ]
        date_matches = [p for p in date_patterns if re.search(p, full_text)]
        if len(date_matches) > 0:
            print(f"PASS: Component 2 — Date found in document (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — No date found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Recipient address block (0.15 points)
    try:
        addr_checks = 0
        if "michael chen" in full_text:
            addr_checks += 1
        if "456 oak avenue" in full_text:
            addr_checks += 1
        if "portland" in full_text and "97201" in full_text:
            addr_checks += 1

        if addr_checks >= 3:
            print(f"PASS: Component 3 — Full recipient address found: Michael Chen, 456 Oak Avenue, Portland, OR 97201 (0.15 pts)")
            total_score += 0.15
        elif addr_checks >= 2:
            print(f"PARTIAL: Component 3 — Partial address found ({addr_checks}/3 parts) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — Recipient address not found ({addr_checks}/3 parts)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Greeting (0.10 points)
    try:
        has_greeting = (
            "dear mr. chen" in full_text
            or "dear mr chen" in full_text
            or "dear michael" in full_text
        )
        if has_greeting:
            print(f"PASS: Component 4 — Proper greeting found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — No greeting 'Dear Mr. Chen' or 'Dear Michael' found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Offer details — position, salary, start date (0.25 points)
    try:
        offer_checks = 0
        if "senior software engineer" in full_text:
            offer_checks += 1
        # Check salary: $145,000 or 145000 or 145,000
        if "145,000" in full_text or "145000" in full_text:
            offer_checks += 1
        # Check start date: May 15, 2026
        if "may 15" in full_text and "2026" in full_text:
            offer_checks += 1

        if offer_checks >= 3:
            print(f"PASS: Component 5 — All offer details found: position, salary, start date (0.25 pts)")
            total_score += 0.25
        elif offer_checks >= 2:
            print(f"PARTIAL: Component 5 — {offer_checks}/3 offer details found (0.15 pts)")
            total_score += 0.15
        elif offer_checks >= 1:
            print(f"PARTIAL: Component 5 — {offer_checks}/3 offer details found (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 5 — No offer details found (position/salary/start date)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Closing with HR Director signature block (0.25 points)
    try:
        closing_checks = 0
        # Check for closing phrase like "sincerely" or "regards"
        if re.search(r'\b(sincerely|regards|respectfully|best regards)\b', full_text):
            closing_checks += 1
        # Check for HR Director title
        if "hr director" in full_text:
            closing_checks += 1

        if closing_checks >= 2:
            print(f"PASS: Component 6 — Closing with HR Director signature block found (0.25 pts)")
            total_score += 0.25
        elif closing_checks >= 1:
            print(f"PARTIAL: Component 6 — Partial closing found ({closing_checks}/2 parts) (0.12 pts)")
            total_score += 0.12
        else:
            print(f"FAIL: Component 6 — No closing or HR Director signature block found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
