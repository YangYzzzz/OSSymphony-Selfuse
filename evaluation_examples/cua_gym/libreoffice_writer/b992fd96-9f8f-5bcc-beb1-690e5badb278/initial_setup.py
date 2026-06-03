"""
Initial Setup: Create a book draft document with 5 Heading 1 paragraphs, no chapter numbering.
Task ID: writer_tm_061
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_061'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title page
    title = doc.add_heading('Book Draft: Data-Driven Decision Making in Modern Organizations', level=0)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(24)
    run = subtitle.add_run('A Comprehensive Study on Analytics-Based Management Practices')
    run.italic = True
    run.font.size = Pt(14)

    author = doc.add_paragraph()
    author.paragraph_format.space_before = Pt(48)
    run = author.add_run('Dr. Elena Vasquez')
    run.font.size = Pt(12)
    author.add_run('\n')
    run2 = author.add_run('Department of Business Analytics, Stanford University')
    run2.font.size = Pt(11)
    run2.italic = True

    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'The landscape of organizational decision-making has undergone a fundamental '
        'transformation over the past two decades. With the exponential growth of digital '
        'data and the maturation of analytical tools, businesses across industries have '
        'increasingly turned to data-driven methodologies to guide strategic and operational choices.'
    )
    doc.add_paragraph(
        'This study examines how organizations at various stages of analytical maturity '
        'leverage data to improve outcomes. We analyze 347 firms across 12 industry sectors, '
        'spanning from early-stage startups to Fortune 500 corporations, over a five-year period '
        'from 2019 to 2024.'
    )
    doc.add_paragraph(
        'Our research questions focus on three key areas: (1) What factors determine the '
        'successful adoption of data-driven practices? (2) How does analytical maturity '
        'correlate with measurable business outcomes? (3) What organizational structures '
        'best support a culture of evidence-based decision making?'
    )

    # Chapter 2: Background
    doc.add_heading('Background', level=1)

    doc.add_paragraph(
        'The concept of data-driven decision making (DDDM) emerged from the broader field '
        'of business intelligence in the early 2000s. Pioneering work by Davenport and Harris '
        '(2007) established the theoretical framework for competing on analytics, arguing that '
        'organizations could gain sustainable competitive advantage through superior use of data.'
    )
    doc.add_paragraph(
        'Subsequent research by Brynjolfsson, Hitt, and Kim (2011) provided empirical evidence '
        'linking DDDM adoption to a 5-6% improvement in productivity and output. Their study of '
        '179 large publicly traded firms demonstrated that the relationship held even after '
        'controlling for traditional measures of IT investment.'
    )
    doc.add_paragraph(
        'More recent scholarship has expanded the focus to include organizational culture, '
        'talent management, and governance structures as critical enablers. McAfee and '
        'Brynjolfsson (2012) introduced the concept of "big data management" capabilities, '
        'emphasizing that technology alone is insufficient without complementary organizational changes.'
    )
    doc.add_paragraph(
        'The evolution of cloud computing platforms, machine learning frameworks, and self-service '
        'analytics tools has dramatically lowered the barriers to entry. Organizations that previously '
        'lacked the resources for sophisticated analysis can now deploy advanced predictive models '
        'with minimal infrastructure investment.'
    )

    # Chapter 3: Methodology
    doc.add_heading('Methodology', level=1)

    doc.add_paragraph(
        'Our research employs a mixed-methods approach combining quantitative survey data with '
        'qualitative case studies. The primary data collection instrument was a 78-item questionnaire '
        'administered to senior analytics leaders and C-suite executives at 347 organizations.'
    )
    doc.add_paragraph(
        'The survey instrument was developed through an iterative process involving pilot testing '
        'with 42 practitioners and academic review by experts in organizational behavior and '
        'information systems. Cronbach\'s alpha for the composite scales ranged from 0.82 to 0.94, '
        'indicating strong internal consistency.'
    )
    doc.add_paragraph(
        'Analytical maturity was assessed using a proprietary five-level framework: '
        '(1) Ad Hoc Reporting, (2) Standardized Analytics, (3) Predictive Modeling, '
        '(4) Prescriptive Optimization, and (5) Autonomous Decision Systems. Each level was '
        'operationalized through observable indicators validated by industry practitioners.'
    )
    doc.add_paragraph(
        'Quantitative analysis included hierarchical linear modeling (HLM) to account for '
        'industry-level variance, structural equation modeling (SEM) for testing mediation '
        'hypotheses, and propensity score matching to address selection bias in the relationship '
        'between analytics investment and business outcomes.'
    )

    # Chapter 4: Results
    doc.add_heading('Results', level=1)

    doc.add_paragraph(
        'Our analysis reveals several significant findings across the three research domains. '
        'First, organizational culture emerged as the strongest predictor of successful DDDM '
        'adoption, explaining 34% of the variance in our analytical maturity index (AMI). '
        'Technology investment, while necessary, accounted for only 18% of the variance.'
    )
    doc.add_paragraph(
        'Second, firms at Level 4 (Prescriptive Optimization) or above demonstrated a mean '
        'revenue growth rate of 12.3% compared to 7.1% for firms at Levels 1-2. After controlling '
        'for industry effects, firm size, and prior performance, the adjusted difference remained '
        'statistically significant at 4.8 percentage points (p < 0.001).'
    )
    doc.add_paragraph(
        'The structural equation model confirmed that data literacy among middle management serves '
        'as a critical mediating variable. Organizations investing in broad-based analytics training '
        'programs showed 2.7 times greater improvement in their AMI scores over the study period '
        'compared to those focusing exclusively on specialist hiring.'
    )
    doc.add_paragraph(
        'Cross-industry analysis revealed notable heterogeneity. Financial services and technology '
        'sectors showed the highest baseline maturity levels, while manufacturing and healthcare '
        'demonstrated the most rapid improvement trajectories during the study period.'
    )

    # Chapter 5: Conclusion
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        'This study contributes to the growing body of evidence supporting the strategic value '
        'of data-driven decision making. Our findings suggest that organizations seeking to advance '
        'their analytical capabilities should prioritize cultural transformation and broad-based '
        'data literacy over pure technology investment.'
    )
    doc.add_paragraph(
        'The five-level maturity framework provides a practical roadmap for organizational leaders, '
        'while our quantitative results offer compelling evidence for the business case of analytics '
        'investment. The 4.8 percentage point growth differential between high and low maturity '
        'firms translates to substantial economic value over multi-year horizons.'
    )
    doc.add_paragraph(
        'Future research should extend this work to examine the role of artificial intelligence '
        'and automated decision systems, which are rapidly reshaping the analytical landscape. '
        'Additionally, longitudinal studies tracking individual firms through maturity transitions '
        'would provide valuable insights into the dynamics of organizational change.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
