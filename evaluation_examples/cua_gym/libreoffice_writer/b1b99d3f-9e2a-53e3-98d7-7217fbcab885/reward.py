"""
Reward Script: Multi-page Employee Evaluation Form
Task ID: writer_hr_044
Domain: libreoffice_writer
Scoring:
  Component 1: Page breaks (3 expected) — 0.20 points
  Component 2: Employee info fields on page 1 — 0.20 points
  Component 3: Performance metrics table (8 categories + header) — 0.25 points
  Component 4: Goals tables (Previous + New) — 0.20 points
  Component 5: Signatures section (employee + manager) — 0.15 points
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_044'


def count_page_breaks(doc):
    """Count manual page breaks inside runs."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                if br.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                    count += 1
        # Also count page_break_before property
        if para.paragraph_format.page_break_before:
            count += 1
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = '\n'.join(p.text for p in doc.paragraphs)

    # Component 1: Page breaks — 3 expected between 4 sections (0.20 points)
    try:
        pb_count = count_page_breaks(doc)
        if pb_count >= 3:
            print(f"PASS: Component 1 — Found {pb_count} page breaks (need >= 3) (0.20 pts)")
            total_score += 0.20
        elif pb_count >= 1:
            partial = round(0.20 * (pb_count / 3.0), 2)
            print(f"PARTIAL: Component 1 — Found {pb_count}/3 page breaks ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No page breaks found (expected >= 3)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Employee info fields on page 1 (0.20 points)
    # The task requires: name, ID, department, manager, review period fields
    try:
        required_fields = [
            r'employee\s*name',
            r'employee\s*id',
            r'department',
            r'(direct\s*)?manager',
            r'review\s*period',
        ]
        all_text_lower = all_text.lower()
        found_fields = 0
        for pattern in required_fields:
            if re.search(pattern, all_text_lower):
                found_fields += 1

        if found_fields >= 5:
            print(f"PASS: Component 2 — All 5 employee info fields found (0.20 pts)")
            total_score += 0.20
        elif found_fields >= 3:
            partial = round(0.20 * (found_fields / 5.0), 2)
            print(f"PARTIAL: Component 2 — {found_fields}/5 employee info fields found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {found_fields}/5 employee info fields found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Performance metrics table — 8 categories + header row (0.25 points)
    # Must have a table with at least 9 rows (1 header + 8 categories)
    try:
        perf_table_found = False
        perf_row_count = 0
        for table in doc.tables:
            # Check if this looks like the performance metrics table
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            if any('category' in h or 'rating' in h or 'performance' in h for h in header_cells):
                perf_table_found = True
                perf_row_count = len(table.rows)
                break

        if perf_table_found and perf_row_count >= 9:
            print(f"PASS: Component 3 — Performance metrics table found with {perf_row_count} rows (>= 9) (0.25 pts)")
            total_score += 0.25
        elif perf_table_found and perf_row_count >= 5:
            partial = round(0.25 * ((perf_row_count - 1) / 8.0), 2)
            print(f"PARTIAL: Component 3 — Performance table found but only {perf_row_count} rows ({partial} pts)")
            total_score += partial
        elif perf_table_found:
            print(f"PARTIAL: Component 3 — Performance table found but only {perf_row_count} rows (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — No performance metrics table found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Goals tables — Previous Goals and New Goals (0.20 points)
    # Need 2 tables (in addition to the performance table): one for previous goals, one for new goals
    try:
        goals_tables = 0
        has_previous_goals_table = False
        has_new_goals_table = False
        for table in doc.tables:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            header_text = ' '.join(header_cells)
            # Previous goals table typically has "status" or "met" in header
            if any(kw in header_text for kw in ['met', 'status']):
                has_previous_goals_table = True
                goals_tables += 1
            # New goals table typically has "target" or "criteria" or "deadline" in header
            elif any(kw in header_text for kw in ['target', 'criteria', 'deadline']):
                has_new_goals_table = True
                goals_tables += 1

        # Also check paragraph text for goals section headers
        has_goals_heading = bool(re.search(r'goal', all_text_lower))

        if has_previous_goals_table and has_new_goals_table:
            print(f"PASS: Component 4 — Both Previous Goals and New Goals tables found (0.20 pts)")
            total_score += 0.20
        elif goals_tables >= 1 and has_goals_heading:
            print(f"PARTIAL: Component 4 — {goals_tables} goals table(s) found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Goals tables not found (previous={has_previous_goals_table}, new={has_new_goals_table})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Signatures section — employee and manager signature blocks (0.15 points)
    try:
        has_employee_sig = bool(re.search(r'employee\s*signature', all_text_lower))
        has_manager_sig = bool(re.search(r'manager\s*signature', all_text_lower))
        has_date_line = bool(re.search(r'date\s*:', all_text_lower) or re.search(r'date\s*_', all_text_lower))

        sig_checks = sum([has_employee_sig, has_manager_sig, has_date_line])

        if sig_checks >= 3:
            print(f"PASS: Component 5 — Signatures section complete (employee sig, manager sig, date lines) (0.15 pts)")
            total_score += 0.15
        elif sig_checks >= 2:
            print(f"PARTIAL: Component 5 — {sig_checks}/3 signature elements found (0.10 pts)")
            total_score += 0.10
        elif sig_checks >= 1:
            print(f"PARTIAL: Component 5 — {sig_checks}/3 signature elements found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No signature elements found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
