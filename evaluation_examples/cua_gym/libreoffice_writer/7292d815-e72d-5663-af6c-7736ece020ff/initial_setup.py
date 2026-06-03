"""
Initial Setup: Create employee handbook document with headings but no TOC
Task ID: writer_hr_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_text(doc, text):
    """Add a normal paragraph with body text."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)
    return para


def create_initial():
    doc = Document()

    # === Page Setup ===
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # === Title Page ===
    # Add spacing before title
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_heading('Pinnacle Solutions Inc.', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('Employee Handbook', level=0)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    spacer = doc.add_paragraph('')
    spacer.paragraph_format.space_after = Pt(24)

    info = doc.add_paragraph('Effective Date: January 1, 2025')
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info2 = doc.add_paragraph('Human Resources Department')
    info2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info3 = doc.add_paragraph('Version 4.2 — Confidential')
    info3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Page break after title page
    doc.add_page_break()

    # === Section 1: Introduction and Welcome ===
    doc.add_heading('Introduction and Welcome', level=1)

    doc.add_heading('Welcome Message', level=2)
    add_body_text(doc, 'Welcome to Pinnacle Solutions Inc. We are thrilled to have you as part of our growing team. This handbook has been prepared to help you understand the policies, benefits, and expectations that make our workplace a productive and enjoyable environment. At Pinnacle Solutions, we believe that every team member plays a vital role in our collective success.')
    add_body_text(doc, 'Our company was founded in 2008 by Sarah Mitchell and David Park with the mission of delivering innovative technology solutions to mid-market enterprises. Since then, we have grown to over 1,200 employees across offices in San Francisco, Austin, Chicago, and London.')

    doc.add_heading('Company Mission and Values', level=2)
    add_body_text(doc, 'Our mission is to empower businesses through intelligent technology solutions that drive growth and operational excellence. We are guided by five core values:')
    add_body_text(doc, '1. Innovation — We embrace creative thinking and continuously seek better ways to solve problems for our clients and colleagues.')
    add_body_text(doc, '2. Integrity — We conduct all business with honesty, transparency, and accountability.')
    add_body_text(doc, '3. Collaboration — We believe the best results come from diverse teams working together toward shared goals.')
    add_body_text(doc, '4. Excellence — We hold ourselves to the highest standards of quality in everything we deliver.')
    add_body_text(doc, '5. Respect — We treat every individual with dignity and value the unique perspectives each person brings.')

    doc.add_heading('How to Use This Handbook', level=2)
    add_body_text(doc, 'This handbook is designed to be your primary reference for workplace policies and procedures. It is organized into eight main sections covering employment practices, compensation, benefits, workplace conduct, safety, technology use, professional development, and separation procedures. Please read it thoroughly during your onboarding period and refer back to it whenever you have questions.')
    add_body_text(doc, 'If you need clarification on any policy described herein, please contact the Human Resources department at hr@pinnaclesolutions.com or extension 4500. We encourage open communication and are always happy to assist.')

    # === Section 2: Employment Policies ===
    doc.add_heading('Employment Policies', level=1)

    doc.add_heading('Equal Employment Opportunity', level=2)
    add_body_text(doc, 'Pinnacle Solutions Inc. is an equal opportunity employer. We do not discriminate on the basis of race, color, religion, sex, national origin, age, disability, genetic information, veteran status, sexual orientation, gender identity, or any other legally protected characteristic. This policy applies to all aspects of employment including recruitment, hiring, promotion, compensation, training, and termination.')
    add_body_text(doc, 'Any employee who believes they have experienced or witnessed discrimination should report it immediately to Human Resources or use the anonymous ethics hotline at 1-800-555-0147.')

    doc.add_heading('Employment Classifications', level=2)
    add_body_text(doc, 'Employees at Pinnacle Solutions are classified as follows:')
    add_body_text(doc, 'Full-Time Regular — Employees scheduled to work 40 hours per week on a regular basis. These employees are eligible for the full benefits package described in Section 4 of this handbook.')
    add_body_text(doc, 'Part-Time Regular — Employees scheduled to work fewer than 30 hours per week. Part-time employees may be eligible for limited benefits as determined by their specific employment agreement.')
    add_body_text(doc, 'Temporary/Contract — Employees hired for a specific project or time period. Temporary employees are not eligible for company benefits unless otherwise specified in their contract.')
    add_body_text(doc, 'Intern — Individuals participating in a structured internship program. Interns may be paid or unpaid in accordance with applicable labor laws.')

    doc.add_heading('Probationary Period', level=2)
    add_body_text(doc, 'All new employees undergo a 90-day probationary period beginning on their first day of employment. During this period, your manager will provide regular feedback on your performance and integration with the team. Successful completion of the probationary period is required for continued employment, though employment remains at-will throughout.')
    add_body_text(doc, 'During the probationary period, employees accrue vacation and sick leave but may not use vacation time until the period is completed. Health insurance coverage begins on the first day of the month following your start date.')

    # === Section 3: Compensation and Payroll ===
    doc.add_heading('Compensation and Payroll', level=1)

    doc.add_heading('Pay Schedule and Direct Deposit', level=2)
    add_body_text(doc, 'Employees are paid on a bi-weekly basis, with pay periods ending every other Friday. Paychecks are distributed the following Wednesday. All employees are strongly encouraged to enroll in direct deposit through the HR portal. Paper checks will be mailed to the address on file for employees who do not enroll in direct deposit.')
    add_body_text(doc, 'Pay stubs are available electronically through the Pinnacle HR Portal at hrportal.pinnaclesolutions.com. If you notice any discrepancies in your pay, please contact Payroll at payroll@pinnaclesolutions.com within 30 days of the pay period in question.')

    doc.add_heading('Overtime Policy', level=2)
    add_body_text(doc, 'Non-exempt employees are entitled to overtime pay at a rate of 1.5 times their regular hourly rate for all hours worked in excess of 40 hours in a single workweek. All overtime must be approved in advance by your direct manager. Unauthorized overtime may result in disciplinary action.')
    add_body_text(doc, 'Exempt employees are expected to work the hours necessary to complete their job responsibilities and are not eligible for overtime compensation. Managers should monitor workloads to ensure reasonable expectations are maintained.')

    doc.add_heading('Performance Reviews and Merit Increases', level=2)
    add_body_text(doc, 'Performance reviews are conducted annually during the first quarter of each fiscal year. Reviews assess accomplishments from the prior year and establish goals for the coming year. Merit increases, when budgeted, are effective April 1 and are based on individual performance ratings:')
    add_body_text(doc, 'Exceptional (5) — 5-7% increase. Exceeds Expectations (4) — 3-5% increase. Meets Expectations (3) — 1-3% increase. Needs Improvement (2) — 0% increase with Performance Improvement Plan. Unsatisfactory (1) — 0% increase with potential termination review.')

    # === Section 4: Benefits ===
    doc.add_heading('Benefits', level=1)

    doc.add_heading('Health Insurance', level=2)
    add_body_text(doc, 'Pinnacle Solutions offers comprehensive medical, dental, and vision insurance to all full-time employees and their eligible dependents. The company subsidizes 80% of employee premiums and 60% of dependent premiums. Three plan options are available:')
    add_body_text(doc, 'PPO Gold Plan — Lowest deductible ($500 individual / $1,000 family), widest provider network. Employee monthly premium: $125.')
    add_body_text(doc, 'PPO Silver Plan — Moderate deductible ($1,500 individual / $3,000 family), broad provider network. Employee monthly premium: $85.')
    add_body_text(doc, 'HDHP Bronze Plan — High deductible ($3,000 individual / $6,000 family) paired with Health Savings Account. Employee monthly premium: $45. Company contributes $750 annually to HSA.')

    doc.add_heading('Retirement Plan', level=2)
    add_body_text(doc, 'The company offers a 401(k) retirement savings plan through Vanguard. Employees may contribute up to the annual IRS maximum. Pinnacle Solutions matches employee contributions dollar-for-dollar up to 4% of base salary, with immediate vesting. Employees are automatically enrolled at a 3% contribution rate upon hire unless they opt out or change their rate within 30 days.')
    add_body_text(doc, 'In addition to the 401(k), the company provides a discretionary profit-sharing contribution based on annual company performance, typically ranging from 2-5% of eligible compensation.')

    doc.add_heading('Paid Time Off', level=2)
    add_body_text(doc, 'Pinnacle Solutions provides a combined Paid Time Off (PTO) bank rather than separate vacation and personal day allocations. PTO accrual rates are based on years of service:')
    add_body_text(doc, '0-2 years: 15 days per year (accrued at 4.62 hours per bi-weekly pay period). 3-5 years: 20 days per year (accrued at 6.15 hours per bi-weekly pay period). 6-10 years: 25 days per year (accrued at 7.69 hours per bi-weekly pay period). 11+ years: 30 days per year (accrued at 9.23 hours per bi-weekly pay period).')
    add_body_text(doc, 'PTO may be carried over up to a maximum of 40 hours into the next calendar year. Any balance exceeding 40 hours on December 31 will be forfeited. Employees separating from the company will be paid for unused PTO up to a maximum of 80 hours.')

    # === Section 5: Workplace Conduct ===
    doc.add_heading('Workplace Conduct', level=1)

    doc.add_heading('Code of Professional Conduct', level=2)
    add_body_text(doc, 'All employees are expected to conduct themselves professionally at all times, whether in the office, at client sites, at company events, or when representing Pinnacle Solutions in any capacity. Professional conduct includes but is not limited to: treating colleagues and clients with courtesy and respect, maintaining a clean and organized workspace, adhering to scheduled work hours, and communicating openly and constructively.')
    add_body_text(doc, 'Employees are representatives of Pinnacle Solutions and should be mindful that their actions, both in person and online, can reflect on the company.')

    doc.add_heading('Anti-Harassment Policy', level=2)
    add_body_text(doc, 'Pinnacle Solutions is committed to providing a workplace free from harassment of any kind. Harassment includes unwelcome verbal, physical, or visual conduct that creates an intimidating, hostile, or offensive work environment. This includes but is not limited to harassment based on sex, race, religion, national origin, disability, age, or any other protected characteristic.')
    add_body_text(doc, 'Any employee who experiences or witnesses harassment should immediately report it to their manager, Human Resources, or the anonymous ethics hotline. All complaints will be investigated promptly and confidentially. Retaliation against any employee who reports harassment in good faith is strictly prohibited and will result in disciplinary action up to and including termination.')

    doc.add_heading('Attendance and Punctuality', level=2)
    add_body_text(doc, 'Regular attendance and punctuality are essential to the smooth operation of our business. Employees are expected to report to work on time and ready to perform their duties at the start of their scheduled shift. If you will be late or absent, you must notify your direct manager as soon as possible, and no later than 30 minutes before your scheduled start time.')
    add_body_text(doc, 'Excessive absenteeism or tardiness may result in disciplinary action. Three or more consecutive days of absence without notification will be considered job abandonment and may result in immediate termination.')

    # === Section 6: Health and Safety ===
    doc.add_heading('Health and Safety', level=1)

    doc.add_heading('Workplace Safety Standards', level=2)
    add_body_text(doc, 'Pinnacle Solutions is committed to providing a safe and healthy workplace for all employees. The company complies with all applicable federal, state, and local health and safety regulations, including OSHA requirements. All employees share responsibility for maintaining a safe work environment.')
    add_body_text(doc, 'Safety Data Sheets (SDS) for all hazardous materials used on premises are maintained in the Facilities office on the first floor and are available electronically through the HR portal. Emergency exits, fire extinguishers, and first aid kits are located on every floor as indicated by posted maps.')

    doc.add_heading('Emergency Procedures', level=2)
    add_body_text(doc, 'In the event of a fire, earthquake, or other emergency, employees should follow the established evacuation procedures posted in all common areas. Assembly points are located in the main parking lot (Lot A) and the adjacent park for overflow. Floor wardens are designated on each floor and are identified by orange vests.')
    add_body_text(doc, 'Fire drills are conducted quarterly. All employees must participate. In the event of a medical emergency, call 911 first, then notify the front desk at extension 0. Trained first responders are available on each floor during business hours.')

    doc.add_heading('Workers Compensation', level=2)
    add_body_text(doc, 'All employees are covered by workers compensation insurance for injuries or illnesses arising out of and in the course of employment. If you are injured on the job, report the injury to your supervisor and Human Resources immediately, regardless of severity. Prompt reporting ensures timely access to medical care and benefits.')
    add_body_text(doc, 'Workers compensation benefits include coverage of medical expenses, temporary disability payments, and rehabilitation services. Pinnacle Solutions will not retaliate against any employee who files a workers compensation claim in good faith.')

    # === Section 7: Technology and Information Security ===
    doc.add_heading('Technology and Information Security', level=1)

    doc.add_heading('Acceptable Use of Technology', level=2)
    add_body_text(doc, 'Company-provided technology resources, including computers, mobile devices, email, internet access, and software, are provided primarily for business use. Limited personal use is permitted provided it does not interfere with work responsibilities, consume excessive bandwidth, or violate any company policy.')
    add_body_text(doc, 'Prohibited uses include: accessing or distributing inappropriate or offensive content, installing unauthorized software, using company resources for personal business ventures, sharing login credentials, and circumventing security controls.')

    doc.add_heading('Data Protection and Confidentiality', level=2)
    add_body_text(doc, 'All employees have a responsibility to protect Pinnacle Solutions confidential information and the personal data of our clients, partners, and colleagues. Confidential information includes trade secrets, financial data, client lists, product roadmaps, employee records, and any information marked as confidential or proprietary.')
    add_body_text(doc, 'Employees must: use strong, unique passwords changed every 90 days; lock workstations when unattended; encrypt sensitive files before transmission; use approved file-sharing platforms only; and report any suspected data breach to IT Security at security@pinnaclesolutions.com immediately.')

    doc.add_heading('Remote Work Technology Requirements', level=2)
    add_body_text(doc, 'Employees approved for remote work must maintain a secure and reliable technology environment. Requirements include: a company-issued laptop with current security patches, a stable internet connection with minimum 25 Mbps download speed, use of the company VPN for all work activities, and a private workspace that protects confidential information from unauthorized viewing or listening.')
    add_body_text(doc, 'The IT department provides technical support for remote workers during standard business hours (8:00 AM - 6:00 PM local time). For urgent issues outside these hours, contact the after-hours support line at 1-800-555-0199.')

    # === Section 8: Separation and Offboarding ===
    doc.add_heading('Separation and Offboarding', level=1)

    doc.add_heading('Voluntary Resignation', level=2)
    add_body_text(doc, 'Employees who wish to voluntarily resign from their position should submit a written resignation to their manager and Human Resources at least two weeks prior to their intended last day of work. Directors and above are requested to provide 30 days notice when possible. The company appreciates extended notice periods as they allow for smooth transitions.')
    add_body_text(doc, 'Resigning employees will participate in an exit interview conducted by Human Resources. The exit interview provides an opportunity to discuss your experience and offer feedback that can help improve the workplace for future employees.')

    doc.add_heading('Return of Company Property', level=2)
    add_body_text(doc, 'Upon separation from the company, employees must return all company-owned property on or before their last day of work. This includes but is not limited to: laptop computers and peripherals, mobile devices, access badges and keys, company credit cards, parking passes, and any confidential documents or materials. Failure to return company property may result in deductions from the final paycheck to the extent permitted by law.')

    doc.add_heading('Final Pay and Benefits Continuation', level=2)
    add_body_text(doc, 'Final paychecks, including any accrued but unused PTO (up to 80 hours), will be processed in accordance with state law. In most cases, final pay is issued on the next regular pay date. Employees who are involuntarily terminated will receive their final paycheck within the timeframe required by applicable state law.')
    add_body_text(doc, 'Health insurance coverage continues through the last day of the month in which separation occurs. Eligible employees will receive COBRA continuation coverage information from our benefits administrator within 14 days of separation. COBRA allows you to continue group health coverage at your own expense for up to 18 months.')

    # Add some trailing content for realism
    doc.add_page_break()
    doc.add_heading('Acknowledgment', level=1)
    add_body_text(doc, 'I acknowledge that I have received, read, and understand the Pinnacle Solutions Inc. Employee Handbook. I understand that this handbook is not a contract of employment and that my employment remains at-will. I agree to abide by the policies and procedures described herein.')
    add_body_text(doc, '')
    add_body_text(doc, 'Employee Signature: ____________________________    Date: ____________')
    add_body_text(doc, '')
    add_body_text(doc, 'Printed Name: ________________________________')
    add_body_text(doc, '')
    add_body_text(doc, 'HR Representative: ____________________________    Date: ____________')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
