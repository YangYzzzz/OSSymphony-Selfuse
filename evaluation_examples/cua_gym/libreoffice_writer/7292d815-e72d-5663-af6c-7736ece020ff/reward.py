"""
Reward Script: Insert Table of Contents in Employee Handbook
Task ID: writer_hr_026
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.25): "Table of Contents" heading exists as Heading 1
  - Component 2 (0.35): TOC field code (instrText with "TOC") exists
  - Component 3 (0.20): TOC is positioned before first content section heading
  - Component 4 (0.20): Original document content is preserved (heading count intact)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_026'


def persist_app_state(domain: str):
    """Attempt to save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: "Table of Contents" heading exists as Heading 1 (0.25 points)
    # The golden file adds a Heading 1 titled "Table of Contents" that doesn't exist in initial
    try:
        toc_heading_index = next(
            (i for i, para in enumerate(doc.paragraphs)
             if para.style.name == 'Heading 1' and 'table of contents' in para.text.lower()),
            -1
        )
        if toc_heading_index >= 0:
            print(f"PASS: Component 1 — 'Table of Contents' heading found at paragraph {toc_heading_index} (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 1 — No 'Table of Contents' Heading 1 found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: TOC field code exists in the document (0.35 points)
    # The golden file has instrText with "TOC" field code; initial has none
    try:
        body = doc.element.body
        instr_texts = body.findall('.//w:instrText', ns)
        toc_instr = next(
            (it for it in instr_texts if it.text and 'TOC' in it.text.upper()),
            None
        )
        if toc_instr is not None:
            print(f"PASS: Component 2 — TOC field code found: {toc_instr.text.strip()!r} (0.35 pts)")
            total_score += 0.35
        else:
            print("FAIL: Component 2 — No TOC field code (instrText with 'TOC') found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: TOC is positioned before the first content section "Introduction and Welcome" (0.20 points)
    # In golden, the TOC heading + content appear before "Introduction and Welcome"
    try:
        intro_index = -1
        toc_heading_idx = -1
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == 'Heading 1':
                if 'table of contents' in para.text.lower():
                    toc_heading_idx = i
                elif 'introduction' in para.text.lower():
                    intro_index = i
                    break
                elif toc_heading_idx == -1:
                    # First non-TOC heading 1 reached without finding TOC first
                    intro_index = i
                    break

        if toc_heading_idx >= 0 and intro_index > toc_heading_idx:
            print(f"PASS: Component 3 — TOC heading (P[{toc_heading_idx}]) is before Introduction (P[{intro_index}]) (0.20 pts)")
            total_score += 0.20
        elif toc_heading_idx == -1:
            print("FAIL: Component 3 — No TOC heading found, cannot verify position")
        else:
            print(f"FAIL: Component 3 — TOC heading (P[{toc_heading_idx}]) is NOT before Introduction (P[{intro_index}])")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Original document content preserved (0.20 points)
    # The initial doc has 9 Heading 1 sections (excluding any TOC heading) and 24 Heading 2 subsections.
    # Golden should have all of these plus the TOC heading.
    try:
        h1_headings = []
        h2_count = 0
        for para in doc.paragraphs:
            if para.style.name == 'Heading 1':
                h1_headings.append(para.text)
            elif para.style.name == 'Heading 2':
                h2_count += 1

        # Filter out the TOC heading from count
        content_h1 = [h for h in h1_headings if 'table of contents' not in h.lower()]
        content_h1_count = len(content_h1)

        # Check that all 9 original H1 sections exist AND there is at least 1 heading that IS a TOC heading
        # (to verify it was added, not that an existing heading was renamed)
        has_toc_heading = any('table of contents' in h.lower() for h in h1_headings)
        h1_ok = content_h1_count >= 8  # at least 8 of the 9 original H1s still present
        h2_ok = h2_count >= 20  # at least 20 of the 24 H2s still present (some tolerance)
        content_added = has_toc_heading and len(h1_headings) > content_h1_count  # TOC was ADDED, not replacing

        if h1_ok and h2_ok and content_added:
            print(f"PASS: Component 4 — Content preserved: {content_h1_count} content H1 headings, {h2_count} H2 headings, TOC added as extra (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — Content integrity: {content_h1_count} H1 (expected >=8), {h2_count} H2 (expected >=20), TOC added: {content_added}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
