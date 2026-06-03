"""
Initial Setup: Create thesis document with numbered citation style
Task ID: writer_bs_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Style setup ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- Title ---
    title = doc.add_heading('Deep Learning Approaches for Natural Language Understanding', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author info ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Elena Marchetti')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil_para.add_run('Department of Computer Science, Stanford University')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph()
    abstract.paragraph_format.first_line_indent = Inches(0.5)
    run = abstract.add_run(
        'This thesis investigates the application of deep learning methodologies '
        'to natural language understanding tasks. We build upon the foundational work '
        'in deep generative models [1] and explore how attention-based architectures '
        'have transformed the field. Our approach leverages recent advances in '
        'pre-trained language representations to achieve state-of-the-art results '
        'on multiple benchmark datasets.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # --- Chapter 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)

    intro_p1 = doc.add_paragraph()
    intro_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = intro_p1.add_run(
        'The rapid advancement of deep learning has fundamentally reshaped the landscape '
        'of artificial intelligence research. Since the publication of the seminal work on '
        'deep learning [1], researchers have developed increasingly sophisticated architectures '
        'capable of learning complex representations from raw data. These developments have '
        'had a particularly profound impact on natural language processing, where traditional '
        'feature engineering approaches have been largely supplanted by end-to-end learned systems.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    intro_p2 = doc.add_paragraph()
    intro_p2.paragraph_format.first_line_indent = Inches(0.5)
    run = intro_p2.add_run(
        'The introduction of the Transformer architecture [2] marked a paradigm shift in '
        'sequence modeling. Unlike recurrent neural networks, which process tokens sequentially, '
        'Transformers employ self-attention mechanisms that allow each token to attend to all '
        'other tokens in the input simultaneously. This parallelizable design not only enables '
        'more efficient training on modern hardware but also captures long-range dependencies '
        'more effectively than previous approaches.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # --- Chapter 2: Related Work ---
    doc.add_heading('2. Related Work', level=1)

    rw_p1 = doc.add_paragraph()
    rw_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = rw_p1.add_run(
        'Residual learning [3] addressed the degradation problem observed in very deep networks, '
        'enabling the training of architectures with hundreds of layers. The skip connections '
        'introduced by this framework have become a ubiquitous component in modern neural network '
        'design, appearing in everything from image classification to language generation models.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    rw_p2 = doc.add_paragraph()
    rw_p2.paragraph_format.first_line_indent = Inches(0.5)
    run = rw_p2.add_run(
        'More recently, the development of pre-trained language models such as BERT [4] has '
        'demonstrated that large-scale unsupervised pre-training followed by task-specific '
        'fine-tuning can achieve remarkable performance across a wide range of NLP benchmarks. '
        'This transfer learning paradigm has become the dominant approach in the field, with '
        'subsequent models like GPT-2 and RoBERTa building upon these foundational ideas.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # --- Chapter 3: Methodology ---
    doc.add_heading('3. Methodology', level=1)

    meth_p1 = doc.add_paragraph()
    meth_p1.paragraph_format.first_line_indent = Inches(0.5)
    run = meth_p1.add_run(
        'Our approach combines the representational power of deep learning [1] with the '
        'attention mechanisms introduced by the Transformer [2]. We propose a novel '
        'architecture that integrates residual connections [3] at multiple scales, drawing '
        'inspiration from both the original residual learning framework and the multi-head '
        'attention mechanism. The model is initialized with pre-trained BERT weights [4] '
        'and further fine-tuned on our target tasks.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    meth_p2 = doc.add_paragraph()
    meth_p2.paragraph_format.first_line_indent = Inches(0.5)
    run = meth_p2.add_run(
        'The training procedure follows a two-stage process. In the first stage, the model '
        'undergoes domain-adaptive pre-training on a large corpus of scientific text. In the '
        'second stage, we fine-tune on task-specific labeled data using a combination of '
        'cross-entropy loss and a novel auxiliary objective that encourages the model to '
        'maintain coherent representations across different granularities of text.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # --- Bibliography (numbered style) ---
    doc.add_page_break()
    doc.add_heading('References', level=1)

    refs = [
        '[1] Goodfellow, I., Bengio, Y., and Courville, A. Deep Learning. MIT Press, 2016.',
        '[2] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., and Polosukhin, I. "Attention Is All You Need." Advances in Neural Information Processing Systems, 2017.',
        '[3] He, K., Zhang, X., Ren, S., and Sun, J. "Deep Residual Learning for Image Recognition." Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, 2016.',
        '[4] Devlin, J., Chang, M., Lee, K., and Toutanova, K. "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding." Proceedings of NAACL-HLT, 2019.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
