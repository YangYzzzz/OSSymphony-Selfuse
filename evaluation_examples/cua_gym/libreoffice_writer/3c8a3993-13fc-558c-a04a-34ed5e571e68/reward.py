"""
Reward Script: Change bibliography citation format to author-year style
Task ID: writer_bs_006
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): In-text citations use author-year format (no numbered refs)
  Component 2 (0.3): Bibliography entries use author-year format (no numbered prefixes, year in parens)
  Component 3 (0.3): Bibliography entries are alphabetically sorted by first author surname
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_006'


def persist_app_state(domain):
    """Try to save any unsaved changes in LibreOffice Writer."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Identify body paragraphs (before References) and reference paragraphs (after References)
    ref_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'References' and p.style and 'Heading' in p.style.name:
            ref_heading_idx = i
            break

    if ref_heading_idx is None:
        print("FAIL: Could not find 'References' heading")
        print("REWARD: 0.0")
        return 0.0

    body_paras = doc.paragraphs[:ref_heading_idx]
    ref_paras = [p for p in doc.paragraphs[ref_heading_idx + 1:] if p.text.strip()]

    body_text = ' '.join(p.text for p in body_paras)

    # Component 1: In-text citations use author-year format (0.4 points)
    # In initial: citations are [1], [2], [3], [4]
    # In golden: citations are [Author, Year] e.g. [Goodfellow, 2016]
    try:
        # Check for numbered citations (should be ABSENT in golden)
        numbered_citations = re.findall(r'\[\d+\]', body_text)
        # Check for author-year citations (should be PRESENT in golden)
        # Pattern: [Surname, YYYY]
        author_year_citations = re.findall(r'\[[A-Z][a-z]+,\s*\d{4}\]', body_text)

        has_no_numbered = len(numbered_citations) == 0
        has_author_year = len(author_year_citations) >= 4  # expect at least 4 citations

        if has_no_numbered and has_author_year:
            print(f"PASS: Component 1 — In-text citations are author-year format. "
                  f"Found {len(author_year_citations)} author-year refs, 0 numbered refs. (0.4 pts)")
            total_score += 0.4
        elif has_author_year and not has_no_numbered:
            # Partial: some converted, some not
            print(f"PARTIAL: Component 1 — Mixed citation styles. "
                  f"{len(author_year_citations)} author-year, {len(numbered_citations)} numbered.")
            total_score += 0.2
        else:
            print(f"FAIL: Component 1 — Expected author-year citations, found "
                  f"{len(numbered_citations)} numbered, {len(author_year_citations)} author-year")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Bibliography entries use author-year format (0.3 points)
    # In initial: entries start with [1], [2], etc. and have "Author. Title. Year."
    # In golden: entries start with "Author (Year). Title." — no numbered prefix
    try:
        if len(ref_paras) == 0:
            print("FAIL: Component 2 — No reference entries found")
        else:
            # Check that NO entries start with numbered prefix like [1], [2]
            has_numbered_prefix = any(re.match(r'^\[\d+\]', p.text.strip()) for p in ref_paras)
            # Check that entries contain year in parentheses like (2016), (2017)
            entries_with_year_parens = sum(
                1 for p in ref_paras
                if re.search(r'\(\d{4}\)', p.text)
            )

            no_numbered = not has_numbered_prefix
            all_have_year = entries_with_year_parens == len(ref_paras)

            if no_numbered and all_have_year:
                print(f"PASS: Component 2 — Bibliography entries use author-year format. "
                      f"{len(ref_paras)} entries, all with (Year), no numbered prefixes. (0.3 pts)")
                total_score += 0.3
            elif no_numbered or all_have_year:
                print(f"PARTIAL: Component 2 — Partially converted. "
                      f"Numbered prefix absent: {no_numbered}, Year in parens: {entries_with_year_parens}/{len(ref_paras)}")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — Entries still in numbered format. "
                      f"Numbered prefix: {has_numbered_prefix}, Year-parens entries: {entries_with_year_parens}/{len(ref_paras)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Bibliography entries are alphabetically sorted by first author surname (0.3 points)
    # In initial: sorted by number [1], [2], [3], [4] (not alphabetical)
    # In golden: sorted alphabetically (Devlin, Goodfellow, He, Vaswani)
    try:
        if len(ref_paras) < 2:
            print("FAIL: Component 3 — Not enough reference entries to check sorting")
        else:
            # Extract first author surname from each entry
            surnames = []
            for p in ref_paras:
                text = p.text.strip()
                # Match surname at start of entry (first word before comma or space)
                m = re.match(r'([A-Za-z\-]+)', text)
                if m:
                    surnames.append(m.group(1))

            if len(surnames) >= 2:
                is_sorted = surnames == sorted(surnames, key=str.lower)
                if is_sorted:
                    print(f"PASS: Component 3 — Bibliography is alphabetically sorted. "
                          f"Order: {surnames} (0.3 pts)")
                    total_score += 0.3
                else:
                    print(f"FAIL: Component 3 — Bibliography not alphabetically sorted. "
                          f"Found: {surnames}, expected: {sorted(surnames, key=str.lower)}")
            else:
                print(f"FAIL: Component 3 — Could not extract enough author surnames. Found: {surnames}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved state before verification
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
