"""
Initial Setup: Format warning box on page 3
Task ID: writer_tech_029
Domain: libreoffice_writer

Creates a multi-page technical document with a plain WARNING paragraph on page 3.
No formatting on the warning paragraph (border, background, bold/red are task goals).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ============ PAGE 1 ============
    h = doc.add_heading('MicroServ Platform - Deployment Guide', level=0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacer

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run('Version 4.2.1 | Last Updated: March 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'MicroServ is a containerized microservices platform designed for '
        'enterprise-grade application deployment. This guide covers the '
        'complete deployment workflow, from initial environment setup '
        'to production rollout and monitoring.'
    )
    doc.add_paragraph(
        'The platform supports multi-region deployments with automatic '
        'failover, horizontal scaling, and integrated observability. '
        'Before proceeding, ensure your infrastructure meets the '
        'minimum requirements outlined in Section 2.'
    )

    doc.add_heading('1.1 Supported Environments', level=2)
    doc.add_paragraph('MicroServ can be deployed on the following platforms:')
    doc.add_paragraph('Amazon Web Services (EKS, ECS, or bare EC2)', style='List Bullet')
    doc.add_paragraph('Google Cloud Platform (GKE or Compute Engine)', style='List Bullet')
    doc.add_paragraph('Microsoft Azure (AKS or Virtual Machines)', style='List Bullet')
    doc.add_paragraph('On-premises Kubernetes clusters (v1.26+)', style='List Bullet')

    doc.add_heading('1.2 Architecture Overview', level=2)
    doc.add_paragraph(
        'The platform consists of five core services: the API Gateway, '
        'the Authentication Service, the Task Scheduler, the Data Pipeline, '
        'and the Monitoring Agent. Each service runs as an independent '
        'container with its own configuration and resource limits.'
    )
    doc.add_paragraph(
        'Inter-service communication uses gRPC for synchronous calls '
        'and Apache Kafka for asynchronous event streaming. All traffic '
        'passes through the API Gateway, which handles rate limiting, '
        'authentication token validation, and request routing.'
    )

    # ============ PAGE 2 ============
    doc.add_page_break()

    doc.add_heading('2. Prerequisites', level=1)
    doc.add_paragraph(
        'Before beginning the deployment process, verify that the '
        'following tools and configurations are in place.'
    )

    doc.add_heading('2.1 Hardware Requirements', level=2)

    # Create a requirements table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Minimum', 'Recommended']
    for i, h_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['CPU Cores', '4 vCPU', '8 vCPU'],
        ['Memory (RAM)', '16 GB', '32 GB'],
        ['Disk Space', '100 GB SSD', '250 GB NVMe SSD'],
        ['Network Bandwidth', '1 Gbps', '10 Gbps'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')

    doc.add_heading('2.2 Software Dependencies', level=2)
    doc.add_paragraph('Docker Engine 24.0 or later', style='List Bullet')
    doc.add_paragraph('Kubernetes CLI (kubectl) v1.28+', style='List Bullet')
    doc.add_paragraph('Helm v3.14+', style='List Bullet')
    doc.add_paragraph('Python 3.11+ (for deployment scripts)', style='List Bullet')
    doc.add_paragraph('PostgreSQL client tools (psql 15+)', style='List Bullet')

    doc.add_heading('2.3 Network Configuration', level=2)
    doc.add_paragraph(
        'Ensure the following ports are open in your firewall configuration:'
    )
    doc.add_paragraph('Port 443 (HTTPS) - API Gateway ingress', style='List Bullet')
    doc.add_paragraph('Port 5432 - PostgreSQL database connections', style='List Bullet')
    doc.add_paragraph('Port 9092 - Kafka broker communication', style='List Bullet')
    doc.add_paragraph('Port 8080 - Internal health check endpoints', style='List Bullet')
    doc.add_paragraph('Port 9090 - Prometheus metrics scraping', style='List Bullet')

    doc.add_paragraph(
        'All external-facing endpoints must be secured with TLS 1.3. '
        'Internal service mesh communication uses mutual TLS (mTLS) '
        'managed by the platform automatically.'
    )

    # ============ PAGE 3 ============
    doc.add_page_break()

    doc.add_heading('3. Configuration Management', level=1)
    doc.add_paragraph(
        'MicroServ uses a centralized configuration system backed by '
        'etcd. Each service reads its configuration at startup and '
        'subscribes to real-time updates via watch keys.'
    )

    doc.add_heading('3.1 Configuration File Structure', level=2)
    doc.add_paragraph(
        'The primary configuration file is located at '
        '/etc/microserv/config.yaml on each node. This file controls '
        'service discovery, logging levels, resource quotas, and '
        'database connection pools.'
    )
    doc.add_paragraph(
        'Configuration changes require a service restart unless '
        'hot-reload is explicitly enabled for the target parameter. '
        'The following parameters support hot-reload: log_level, '
        'rate_limit_threshold, and cache_ttl.'
    )

    # The WARNING paragraph - plain text, no special formatting
    warning_para = doc.add_paragraph()
    warning_run = warning_para.add_run(
        'WARNING: Do not modify the configuration file while the service is running.'
    )
    warning_run.font.size = Pt(11)
    # No bold, no red color, no border, no background -- these are the task goals

    doc.add_paragraph('')

    doc.add_heading('3.2 Environment Variables', level=2)
    doc.add_paragraph(
        'In addition to the YAML configuration, certain runtime '
        'parameters can be overridden using environment variables. '
        'These take precedence over file-based configuration.'
    )
    doc.add_paragraph('MICROSERV_LOG_LEVEL - Controls verbosity (DEBUG, INFO, WARN, ERROR)', style='List Bullet')
    doc.add_paragraph('MICROSERV_DB_HOST - Primary database hostname', style='List Bullet')
    doc.add_paragraph('MICROSERV_CACHE_TTL - Cache time-to-live in seconds', style='List Bullet')
    doc.add_paragraph('MICROSERV_MAX_CONNECTIONS - Connection pool ceiling', style='List Bullet')

    doc.add_heading('3.3 Secret Management', level=2)
    doc.add_paragraph(
        'Sensitive values such as API keys, database passwords, and '
        'TLS certificates must never be stored in the configuration '
        'file. MicroServ integrates with HashiCorp Vault for secret '
        'injection at runtime. Configure the Vault address and '
        'authentication token in the deployment manifest.'
    )

    # ============ PAGE 4 (extra content for realism) ============
    doc.add_page_break()

    doc.add_heading('4. Deployment Procedures', level=1)
    doc.add_paragraph(
        'This section outlines the step-by-step process for deploying '
        'MicroServ to a Kubernetes cluster. Follow these instructions '
        'carefully to avoid service disruption.'
    )

    doc.add_heading('4.1 Initial Deployment', level=2)
    doc.add_paragraph('Clone the deployment repository from the internal GitLab instance.', style='List Number')
    doc.add_paragraph('Copy the sample configuration and update values for your environment.', style='List Number')
    doc.add_paragraph('Run the pre-flight validation script to check cluster readiness.', style='List Number')
    doc.add_paragraph('Execute the Helm install command with the production values file.', style='List Number')
    doc.add_paragraph('Verify all pods reach Ready state within five minutes.', style='List Number')

    doc.add_heading('4.2 Rolling Updates', level=2)
    doc.add_paragraph(
        'For zero-downtime updates, use the rolling update strategy. '
        'The platform automatically drains connections from old pods '
        'before terminating them. Set maxUnavailable to 1 and '
        'maxSurge to 2 for optimal update throughput.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
