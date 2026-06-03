"""
Initial Setup: Compare document versions and selectively accept/reject tracked changes
Task ID: writer_lec_083
Domain: libreoffice_writer

Creates two documents:
  1. writer_lec_083.docx — the current (newer) version with all 4 sections updated
  2. writer_lec_083_old.docx — the older version from "last week" with original text
Opens the current document in LibreOffice Writer.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_083'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
OLD_VERSION = f'{WORKDIR}/{TASK_ID}_old.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# ---------- Shared structure helpers ----------

def add_title(doc):
    """Add a common document title."""
    heading = doc.add_heading("Quarterly Strategic Review — Vanguard Technologies", level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    intro = doc.add_paragraph(
        "This document summarizes key strategic initiatives across four divisions "
        "for Q1 2025. Each section covers objectives, progress metrics, and next steps."
    )
    intro.paragraph_format.space_after = Pt(12)


# ============================================================
# OLD VERSION — represents the document from "last week"
# ============================================================

def create_old_version():
    doc = Document()
    add_title(doc)

    # --- Section 1: Product Development ---
    doc.add_heading("Section 1: Product Development", level=1)
    doc.add_paragraph(
        "The product development team has been focused on the Neptune platform redesign "
        "throughout Q4 2024. The legacy architecture relied on monolithic services that "
        "created bottlenecks during peak traffic periods. Initial benchmarks show a 15% "
        "improvement in response times after migrating the authentication module."
    )
    doc.add_paragraph(
        "Key milestones for this quarter included completing the database migration from "
        "PostgreSQL 12 to PostgreSQL 16, which required careful coordination with the "
        "DevOps team. The migration was executed over a weekend maintenance window with "
        "minimal downtime of approximately 45 minutes."
    )
    doc.add_paragraph(
        "Budget allocation for product development stands at $1.2 million for the current "
        "fiscal year, with 60% earmarked for infrastructure upgrades and 40% for new "
        "feature development. Staffing remains stable at 28 engineers across three pods."
    )

    # --- Section 2: Marketing and Outreach ---
    doc.add_heading("Section 2: Marketing and Outreach", level=1)
    doc.add_paragraph(
        "The marketing division launched the Horizon campaign in November 2024, targeting "
        "mid-market SaaS companies with annual revenue between $5M and $50M. Early metrics "
        "indicated a 3.2% click-through rate on LinkedIn advertisements, which was below "
        "the industry average of 4.1%."
    )
    doc.add_paragraph(
        "Content marketing efforts produced 12 blog posts, 4 whitepapers, and 2 webinars "
        "during the quarter. The webinar on API-first architecture attracted 340 registrants "
        "with a 52% attendance rate. Post-event surveys showed 78% of attendees rated the "
        "content as highly relevant."
    )
    doc.add_paragraph(
        "The total marketing budget for Q1 2025 is $680,000, allocated across digital "
        "advertising ($280,000), content production ($150,000), events ($180,000), and "
        "analytics tools ($70,000)."
    )

    # --- Section 3: Human Resources ---
    doc.add_heading("Section 3: Human Resources", level=1)
    doc.add_paragraph(
        "The HR department conducted the annual employee engagement survey in December 2024. "
        "Participation reached 89% of eligible staff, up from 82% the previous year. Overall "
        "satisfaction scored 7.4 out of 10, with notable improvements in work-life balance "
        "(7.8) and manager effectiveness (7.6)."
    )
    doc.add_paragraph(
        "Recruitment activities yielded 14 new hires across engineering, sales, and customer "
        "success. Time-to-fill averaged 34 days, a slight improvement from the prior "
        "quarter's 38 days. The offer acceptance rate was 85%, consistent with historical norms."
    )
    doc.add_paragraph(
        "Training and development programs consumed $95,000 in Q4, funding 6 external "
        "certifications (AWS, Scrum Master) and 3 internal leadership workshops attended "
        "by 42 managers. The retention rate for the quarter was 96.2%."
    )

    # --- Section 4: Financial Operations ---
    doc.add_heading("Section 4: Financial Operations", level=1)
    doc.add_paragraph(
        "Total revenue for Q4 2024 reached $4.8 million, representing a 6% year-over-year "
        "increase. Subscription revenue accounted for $3.9 million (81%), while professional "
        "services contributed $620,000 and one-time license fees totaled $280,000."
    )
    doc.add_paragraph(
        "Operating expenses were $3.6 million, yielding an operating margin of 25%. The "
        "largest expense categories were personnel ($2.1 million), infrastructure ($720,000), "
        "and marketing ($480,000). Cash reserves at quarter-end stood at $8.2 million."
    )
    doc.add_paragraph(
        "Accounts receivable aging improved, with 92% of invoices collected within 30 days "
        "compared to 87% in Q3. The finance team implemented automated payment reminders "
        "that contributed to a 15% reduction in overdue balances."
    )

    doc.save(OLD_VERSION)
    print(f'Old version created: {OLD_VERSION}')


# ============================================================
# CURRENT VERSION — the updated document (all 4 sections modified)
# ============================================================

def create_current_version():
    doc = Document()
    add_title(doc)

    # --- Section 1: Product Development (UPDATED) ---
    doc.add_heading("Section 1: Product Development", level=1)
    doc.add_paragraph(
        "The product development team has completed the Neptune platform redesign and "
        "successfully launched the microservices architecture in production on January 15, "
        "2025. Performance benchmarks now show a 42% improvement in response times compared "
        "to the legacy monolithic system, exceeding the original target of 30%."
    )
    doc.add_paragraph(
        "Key milestones achieved this quarter include the full database migration to "
        "PostgreSQL 16, deployment of the new CI/CD pipeline using GitHub Actions, and "
        "integration of automated security scanning via Snyk. The migration was completed "
        "ahead of schedule with zero data loss and only 22 minutes of planned downtime."
    )
    doc.add_paragraph(
        "Budget allocation for product development has been increased to $1.5 million for "
        "the current fiscal year, with 50% earmarked for infrastructure upgrades, 35% for "
        "new feature development, and 15% for technical debt reduction. The team has grown "
        "to 32 engineers across four specialized pods."
    )

    # --- Section 2: Marketing and Outreach (UPDATED) ---
    doc.add_heading("Section 2: Marketing and Outreach", level=1)
    doc.add_paragraph(
        "The marketing division expanded the Horizon campaign in January 2025, broadening "
        "the target to include enterprise companies with annual revenue up to $200M. The "
        "revised LinkedIn strategy achieved a 5.7% click-through rate, significantly "
        "surpassing the industry average of 4.1% and the previous quarter's 3.2%."
    )
    doc.add_paragraph(
        "Content marketing efforts accelerated with 18 blog posts, 6 whitepapers, 3 "
        "webinars, and a new podcast series. The webinar on zero-trust security architecture "
        "attracted 580 registrants with a 61% attendance rate. The podcast has accumulated "
        "4,200 downloads in its first month."
    )
    doc.add_paragraph(
        "The total marketing budget for Q1 2025 has been revised upward to $820,000, "
        "allocated across digital advertising ($340,000), content production ($180,000), "
        "events ($200,000), analytics tools ($60,000), and the new podcast initiative "
        "($40,000)."
    )

    # --- Section 3: Human Resources (UPDATED) ---
    doc.add_heading("Section 3: Human Resources", level=1)
    doc.add_paragraph(
        "The HR department completed a follow-up pulse survey in February 2025, achieving "
        "a record 94% participation rate. Overall satisfaction climbed to 8.1 out of 10, "
        "with significant gains in career development opportunities (8.3) and compensation "
        "fairness (7.9). A new mentorship program was launched in response to employee feedback."
    )
    doc.add_paragraph(
        "Recruitment activities have been aggressive, with 23 new hires onboarded across "
        "engineering, product management, sales, and customer success. Time-to-fill decreased "
        "to 27 days thanks to the new applicant tracking system. The offer acceptance rate "
        "rose to 91%, reflecting improved employer branding efforts."
    )
    doc.add_paragraph(
        "Training and development investment increased to $140,000 in Q1, funding 11 "
        "external certifications (AWS, GCP, PMP), 5 internal leadership workshops attended "
        "by 58 managers, and a company-wide AI literacy program. The retention rate for "
        "the quarter improved to 97.8%."
    )

    # --- Section 4: Financial Operations (UPDATED) ---
    doc.add_heading("Section 4: Financial Operations", level=1)
    doc.add_paragraph(
        "Total revenue for Q1 2025 reached $5.6 million, representing an impressive 17% "
        "year-over-year increase. Subscription revenue accounted for $4.7 million (84%), "
        "while professional services contributed $580,000 and one-time license fees totaled "
        "$320,000. Three new enterprise contracts worth a combined $1.1 million were signed."
    )
    doc.add_paragraph(
        "Operating expenses were $4.0 million, yielding an operating margin of 28.6%. The "
        "largest expense categories were personnel ($2.4 million), infrastructure ($780,000), "
        "marketing ($520,000), and R&D ($300,000). Cash reserves at quarter-end grew to "
        "$9.8 million, strengthening the balance sheet."
    )
    doc.add_paragraph(
        "Accounts receivable management continued to improve, with 95% of invoices collected "
        "within 30 days. The finance team deployed a new ERP module for automated invoicing "
        "and reconciliation, reducing manual processing time by 40% and eliminating two "
        "recurring data entry errors."
    )

    doc.save(OUTPUT)
    print(f'Current version created: {OUTPUT}')


# ============================================================
# MAIN
# ============================================================

create_old_version()
create_current_version()

# Open the current document in LibreOffice Writer
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
