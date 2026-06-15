"""
Initial Setup: Create a Writer document with Introduction, Main Content, and Conclusion
all in single-column layout.
Task ID: writer_fs_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default page margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Introduction ---
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'The rapid advancement of renewable energy technologies over the past decade '
        'has fundamentally reshaped the global energy landscape. Solar photovoltaic '
        'installations have seen a remarkable 85% reduction in cost since 2010, making '
        'them competitive with traditional fossil fuel sources in most markets. Wind '
        'energy capacity has similarly expanded, with offshore wind farms now capable '
        'of powering millions of households across coastal regions.'
    )

    doc.add_paragraph(
        'This report examines the current state of renewable energy adoption across '
        'twelve major economies, analyzing investment trends, policy frameworks, and '
        'technological breakthroughs that are driving the transition toward a '
        'sustainable energy future. Our analysis draws on data collected between '
        'January 2024 and December 2025 from government agencies, industry '
        'publications, and independent research institutions.'
    )

    # --- Main Content ---
    doc.add_heading('Main Content', level=1)

    doc.add_paragraph(
        'Solar energy deployment reached unprecedented levels in 2025, with global '
        'installed capacity surpassing 2.1 terawatts. China continued to lead with '
        '412 gigawatts of new installations, followed by the United States at 89 '
        'gigawatts and India at 67 gigawatts. The European Union collectively added '
        '78 gigawatts, driven primarily by aggressive targets set under the Green Deal.'
    )

    doc.add_paragraph(
        'Wind energy experienced significant growth in both onshore and offshore '
        'segments. Denmark\'s Kriegers Flak complex, completed in late 2024, now '
        'generates 1.4 gigawatts of power, enough to supply approximately 1.3 million '
        'homes. The United Kingdom\'s Dogger Bank wind farm, the world\'s largest, '
        'reached full operational capacity in March 2025 at 3.6 gigawatts.'
    )

    doc.add_paragraph(
        'Battery storage technology has emerged as a critical enabler for renewable '
        'energy integration. Lithium iron phosphate batteries now dominate the utility '
        'scale market, with costs declining to $98 per kilowatt-hour. Grid-scale '
        'storage installations totaled 147 gigawatt-hours globally in 2025, a 42% '
        'increase over the previous year.'
    )

    doc.add_paragraph(
        'Hydrogen produced from renewable electricity, commonly termed green hydrogen, '
        'is gaining traction as a decarbonization pathway for heavy industry and '
        'long-distance transport. The European Hydrogen Backbone project has mapped '
        'over 28,000 kilometers of pipeline infrastructure, with construction underway '
        'on the first 4,500-kilometer segment connecting the Netherlands to southern Italy.'
    )

    doc.add_paragraph(
        'Policy frameworks have been instrumental in accelerating adoption. The United '
        'States Inflation Reduction Act allocated $369 billion in clean energy '
        'incentives, catalyzing over $280 billion in private sector investment by the '
        'end of 2025. Japan\'s Green Transformation program committed 20 trillion yen '
        'in transition bonds, while Brazil expanded its auction-based procurement system '
        'to include floating solar and agrivoltaic installations.'
    )

    doc.add_paragraph(
        'Grid modernization remains a significant challenge. Transmission bottlenecks '
        'have delayed the interconnection of 38 gigawatts of approved renewable projects '
        'in the United States alone. The European Network of Transmission System '
        'Operators estimates that $594 billion in grid investment is needed by 2030 to '
        'accommodate planned renewable capacity additions across the continent.'
    )

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        'The data presented in this report demonstrates that the global energy '
        'transition has reached an inflection point. Renewable energy sources now '
        'account for 38% of global electricity generation, up from 29% in 2020. '
        'The economic case for clean energy has become self-reinforcing, with declining '
        'costs driving increased deployment, which in turn attracts greater investment '
        'and further technological improvement.'
    )

    doc.add_paragraph(
        'Looking ahead, the primary barriers to continued progress are no longer '
        'technological or economic but rather institutional and infrastructural. '
        'Streamlining permitting processes, expanding grid capacity, and developing '
        'robust supply chains for critical minerals will be essential to maintaining '
        'the current pace of deployment. International cooperation on technology '
        'transfer and climate finance will be equally important in ensuring that the '
        'benefits of the energy transition are shared equitably across all regions.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
