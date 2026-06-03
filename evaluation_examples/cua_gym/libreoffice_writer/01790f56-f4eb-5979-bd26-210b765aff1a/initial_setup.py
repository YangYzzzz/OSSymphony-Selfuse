"""
Initial Setup: Create a Writer document with multiple heading levels but no numbering.
Task ID: writer_tech_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_058'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Cloud Infrastructure Migration Plan', level=0)

    # --- Section 1: Heading 1 ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines the comprehensive strategy for migrating our '
        'on-premises infrastructure to a cloud-native architecture. The migration '
        'is expected to reduce operational costs by 35% while improving system '
        'reliability and scalability across all business units.'
    )
    doc.add_paragraph(
        'The project timeline spans 18 months, with three major phases covering '
        'assessment, migration execution, and post-migration optimization. Key '
        'stakeholders from Engineering, Operations, and Finance have reviewed and '
        'approved this plan.'
    )

    # --- Section 1.1: Heading 2 ---
    doc.add_heading('Project Scope', level=2)
    doc.add_paragraph(
        'The migration covers 47 production services currently running across '
        'three data centers in Portland, Chicago, and Atlanta. Each service will '
        'be categorized into one of four migration strategies: rehost, replatform, '
        'refactor, or retire.'
    )

    # --- Section 1.1.1: Heading 3 ---
    doc.add_heading('Included Systems', level=3)
    doc.add_paragraph(
        'All customer-facing web applications, internal API gateways, database '
        'clusters (PostgreSQL and MongoDB), message queues (RabbitMQ), and '
        'monitoring infrastructure are included in the migration scope.'
    )

    doc.add_heading('Excluded Systems', level=3)
    doc.add_paragraph(
        'Legacy mainframe applications supporting the payroll system and the '
        'on-premises Active Directory infrastructure will remain in the current '
        'data centers pending a separate modernization initiative scheduled for Q3 2027.'
    )

    # --- Section 1.2: Heading 2 ---
    doc.add_heading('Timeline Overview', level=2)
    doc.add_paragraph(
        'Phase 1 (Months 1-4): Discovery and assessment of all workloads, '
        'dependency mapping, and cost modeling. Phase 2 (Months 5-12): Staged '
        'migration of services in priority order. Phase 3 (Months 13-18): '
        'Optimization, decommissioning of legacy hardware, and knowledge transfer.'
    )

    doc.add_heading('Key Milestones', level=3)
    doc.add_paragraph(
        'Month 2: Complete infrastructure audit and dependency graph. '
        'Month 6: First production workload migrated. '
        'Month 10: 75% of services operational in cloud. '
        'Month 15: All primary services migrated. '
        'Month 18: Legacy hardware decommissioned.'
    )

    # --- Section 2: Heading 1 ---
    doc.add_heading('Technical Architecture', level=1)
    doc.add_paragraph(
        'The target architecture leverages a multi-region deployment model with '
        'primary resources in US-West-2 and failover capacity in US-East-1. '
        'All services will be containerized using Docker and orchestrated via '
        'Kubernetes (EKS) with Istio service mesh for traffic management.'
    )

    # --- Section 2.1: Heading 2 ---
    doc.add_heading('Network Design', level=2)
    doc.add_paragraph(
        'A hub-and-spoke VPC topology will be implemented with dedicated VPCs for '
        'production, staging, and development environments. Transit Gateway will '
        'provide inter-VPC connectivity, and AWS Direct Connect will maintain a '
        '10 Gbps dedicated link to remaining on-premises systems.'
    )

    doc.add_heading('Security Zones', level=3)
    doc.add_paragraph(
        'Network segmentation will enforce three security zones: public-facing '
        '(DMZ), application tier (private subnets), and data tier (isolated subnets '
        'with no direct internet access). Network ACLs and security groups will '
        'implement defense-in-depth at each layer.'
    )

    doc.add_heading('DNS and Load Balancing', level=3)
    doc.add_paragraph(
        'Route 53 will manage DNS with weighted routing for blue-green deployments. '
        'Application Load Balancers will handle HTTPS termination and path-based '
        'routing. Global Accelerator will provide static IP addresses for clients '
        'requiring fixed endpoints.'
    )

    # --- Section 2.2: Heading 2 ---
    doc.add_heading('Data Migration Strategy', level=2)
    doc.add_paragraph(
        'Database migration will use a combination of AWS Database Migration Service '
        'for PostgreSQL workloads and custom ETL pipelines for MongoDB collections. '
        'A dual-write pattern will be employed during the transition period to ensure '
        'zero data loss and enable instant rollback if issues arise.'
    )

    doc.add_heading('Data Validation Framework', level=3)
    doc.add_paragraph(
        'An automated validation framework will compare source and target databases '
        'using row counts, checksum verification, and sample-based deep comparison. '
        'Validation reports will be generated hourly during active migration windows '
        'and daily during steady-state periods.'
    )

    # --- Section 3: Heading 1 ---
    doc.add_heading('Risk Assessment', level=1)
    doc.add_paragraph(
        'A comprehensive risk register has been compiled identifying 23 potential '
        'risks across technical, operational, and organizational categories. Each risk '
        'has been evaluated using a probability-impact matrix and assigned appropriate '
        'mitigation strategies.'
    )

    doc.add_heading('Technical Risks', level=2)
    doc.add_paragraph(
        'The highest-rated technical risks include: data corruption during migration '
        '(mitigated by dual-write and validation), network latency between hybrid '
        'components (mitigated by Direct Connect), and container orchestration '
        'complexity (mitigated by phased rollout and training).'
    )

    doc.add_heading('Latency Sensitivity Analysis', level=3)
    doc.add_paragraph(
        'Services with sub-10ms latency requirements have been identified and will '
        'be co-located within the same availability zone. Cross-region calls will be '
        'minimized through strategic data replication and caching layers using '
        'ElastiCache Redis clusters.'
    )

    doc.add_heading('Operational Risks', level=2)
    doc.add_paragraph(
        'Staff training gaps and potential knowledge silos represent significant '
        'operational risks. A structured training program with AWS certification '
        'targets has been established for all infrastructure team members, with '
        'completion required before Phase 2 begins.'
    )

    # --- Section 4: Heading 1 ---
    doc.add_heading('Budget and Resources', level=1)
    doc.add_paragraph(
        'The total project budget is $2.4 million, covering cloud infrastructure '
        'costs, tooling licenses, professional services, and internal staff allocation. '
        'Monthly cloud spend is projected to stabilize at $185,000 after migration, '
        'compared to the current $285,000 in data center operating costs.'
    )

    doc.add_heading('Cost Breakdown', level=2)
    doc.add_paragraph(
        'Infrastructure provisioning: $890,000. Migration tooling and services: '
        '$420,000. Training and certification: $165,000. Contingency reserve (15%): '
        '$360,000. Internal labor allocation: $565,000.'
    )

    doc.add_heading('Compute Cost Projections', level=3)
    doc.add_paragraph(
        'Reserved instances will cover 70% of baseline compute needs, with savings '
        'plans applied to the remaining steady-state workloads. Spot instances will '
        'handle batch processing and non-critical background jobs, reducing compute '
        'costs by an estimated 60% for those workloads.'
    )

    doc.add_heading('Storage Cost Projections', level=3)
    doc.add_paragraph(
        'Intelligent tiering will automatically move infrequently accessed data to '
        'lower-cost storage classes. Data lifecycle policies will archive objects older '
        'than 90 days to Glacier Deep Archive, with estimated annual storage savings '
        'of $42,000 compared to standard S3 pricing.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
