"""
Reward Script: Change code paragraph background color to light gray (#E8E8E8)
Task ID: writer_tech_017
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Code paragraph has a shading/background element
  Component 2 (0.35): Fill color is exactly E8E8E8
  Component 3 (0.25): Shading is solid (val=clear) and non-code paragraphs unaffected
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_017'


def color_hex_distance(hex1, hex2):
    """Compute RGB distance between two hex color strings (without #)."""
    r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
    r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def find_code_paragraph(doc):
    """Find the paragraph with Liberation Mono font (the code block)."""
    from docx.oxml.ns import qn
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run.font.name and 'mono' in run.font.name.lower():
                return i, para
    return None, None


def get_paragraph_shading(para):
    """Extract shading info from a paragraph's XML properties."""
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        return None
    return {
        'fill': shd.get(qn('w:fill')),
        'val': shd.get(qn('w:val')),
        'color': shd.get(qn('w:color')),
    }


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.oxml.ns import qn

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the code paragraph (Liberation Mono font)
    code_idx, code_para = find_code_paragraph(doc)
    if code_para is None:
        print("FAIL: Could not find a paragraph with monospace font (code block)")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Code paragraph found at index {code_idx}: '{code_para.text[:60]}...'")

    # Component 1: Code paragraph has a shading/background element (0.4 points)
    try:
        shd_info = get_paragraph_shading(code_para)
        if shd_info is not None and shd_info.get('fill') is not None:
            fill_val = shd_info['fill']
            # Exclude 'auto' as that means no explicit fill
            if fill_val.lower() != 'auto':
                print(f"PASS: Component 1 — Code paragraph has shading fill={fill_val} (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 — Shading fill is 'auto' (no explicit background)")
        else:
            print(f"FAIL: Component 1 — No shading element on code paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Fill color is exactly E8E8E8 or very close (0.35 points)
    try:
        shd_info = get_paragraph_shading(code_para)
        if shd_info is not None and shd_info.get('fill') is not None:
            fill_color = shd_info['fill'].upper().replace('#', '')
            target_color = 'E8E8E8'
            dist = color_hex_distance(fill_color, target_color)
            if dist < 5.0:
                print(f"PASS: Component 2 — Fill color {fill_color} matches target {target_color} (distance={dist:.1f}) (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 2 — Fill color {fill_color} does not match target {target_color} (distance={dist:.1f})")
        else:
            print(f"FAIL: Component 2 — No shading to check color")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Shading is solid fill (val=clear) and non-code paragraphs unaffected (0.25 points)
    # NOTE: This component requires shading to exist first (gate on Component 1)
    try:
        shd_info = get_paragraph_shading(code_para)
        has_shading = (shd_info is not None and shd_info.get('fill') is not None
                       and shd_info['fill'].lower() != 'auto')

        if not has_shading:
            # No shading at all — this component cannot pass
            print(f"FAIL: Component 3 — No shading on code paragraph, cannot evaluate quality")
        else:
            solid_fill = False
            val = shd_info.get('val', '')
            if val and val.lower() == 'clear':
                solid_fill = True

            # Check that non-code paragraphs do NOT have unexpected shading
            non_code_shaded = 0
            for i, para in enumerate(doc.paragraphs):
                if i == code_idx:
                    continue
                other_shd = get_paragraph_shading(para)
                if other_shd is not None and other_shd.get('fill') is not None:
                    fill = other_shd['fill'].lower()
                    if fill != 'auto':
                        non_code_shaded += 1

            if solid_fill and non_code_shaded == 0:
                print(f"PASS: Component 3 — Shading is solid (val=clear), no non-code paragraphs shaded (0.25 pts)")
                total_score += 0.25
            elif solid_fill and non_code_shaded > 0:
                print(f"PARTIAL: Component 3 — Shading is solid but {non_code_shaded} other paragraph(s) also shaded (0.1 pts)")
                total_score += 0.1
            elif not solid_fill and non_code_shaded == 0:
                print(f"PARTIAL: Component 3 — Shading val is not 'clear' but no other paragraphs affected (0.1 pts)")
                total_score += 0.1
            else:
                print(f"FAIL: Component 3 — Shading val not 'clear' and {non_code_shaded} other paragraph(s) also shaded")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
