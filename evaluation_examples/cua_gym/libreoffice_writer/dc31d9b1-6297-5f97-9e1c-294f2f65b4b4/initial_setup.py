"""
Initial Setup: Create a Writer document with a code example paragraph in Liberation Mono (no background).
Task ID: writer_tech_017
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_017'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    heading = doc.add_heading('Data Processing Pipeline — Technical Reference', level=1)

    # --- Introductory paragraph ---
    p1 = doc.add_paragraph()
    run1 = p1.add_run(
        'This document describes the core data ingestion module used by the analytics platform. '
        'The pipeline processes incoming CSV files from partner integrations, validates schema '
        'compliance, and loads clean records into the staging database for downstream reporting.'
    )
    run1.font.name = 'Liberation Sans'
    run1.font.size = Pt(11)

    # --- Second paragraph ---
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        'Before running the pipeline, ensure that the PostgreSQL service is active and the '
        'environment variable DB_CONNECTION_STRING is set. The default batch size is 500 rows; '
        'adjust this in config.yaml under the ingestion.batch_size key if memory constraints apply.'
    )
    run2.font.name = 'Liberation Sans'
    run2.font.size = Pt(11)

    # --- Sub-heading ---
    doc.add_heading('Example Usage', level=2)

    # --- Explanation before code ---
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        'The following snippet demonstrates how to invoke the pipeline from a Python script:'
    )
    run3.font.name = 'Liberation Sans'
    run3.font.size = Pt(11)

    # --- Code example paragraph (Liberation Mono, NO background) ---
    code_text = (
        'from pipeline.ingest import DataLoader\n'
        'loader = DataLoader(source="partner_feed", batch_size=500)\n'
        'loader.validate_schema("schemas/partner_v2.json")\n'
        'result = loader.run(dry_run=False)\n'
        'print(f"Loaded {result.rows_inserted} rows in {result.elapsed_sec:.1f}s")'
    )
    p_code = doc.add_paragraph()
    run_code = p_code.add_run(code_text)
    run_code.font.name = 'Liberation Mono'
    run_code.font.size = Pt(10)
    run_code.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # NO background shading — that is what the task asks the agent to add

    # --- Paragraph after code ---
    p4 = doc.add_paragraph()
    run4 = p4.add_run(
        'The dry_run parameter allows testing the full pipeline without committing rows to the '
        'database. When set to True, validation and transformation steps execute normally but the '
        'final INSERT is replaced with a row-count summary written to stdout.'
    )
    run4.font.name = 'Liberation Sans'
    run4.font.size = Pt(11)

    # --- Additional section ---
    doc.add_heading('Error Handling', level=2)

    p5 = doc.add_paragraph()
    run5 = p5.add_run(
        'If schema validation fails, the loader raises a SchemaViolationError with a detailed '
        'report listing each non-conforming column. Partial batches are rolled back automatically. '
        'Retry logic is configurable via config.yaml under ingestion.retry_policy.'
    )
    run5.font.name = 'Liberation Sans'
    run5.font.size = Pt(11)

    p6 = doc.add_paragraph()
    run6 = p6.add_run(
        'For production deployments, enable the dead-letter queue by setting '
        'ingestion.dlq_enabled to true. Failed records are written to the dlq_records table '
        'with the original payload and error metadata for manual review.'
    )
    run6.font.name = 'Liberation Sans'
    run6.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
