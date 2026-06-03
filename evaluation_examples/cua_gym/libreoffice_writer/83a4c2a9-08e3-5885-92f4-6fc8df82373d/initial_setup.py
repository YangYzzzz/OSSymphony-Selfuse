"""
Initial Setup: Set unique line spacing for each of the three body paragraphs
Task ID: osworld_writer_line_spacing_per_paragraph_007
Domain: libreoffice_writer

Creates an academic paper excerpt with 3 body paragraphs, each using 1.15 line spacing.
The agent must change them to: paragraph 1 = 1.0, paragraph 2 = 1.5, paragraph 3 = 2.0.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document title ---
    title = doc.add_heading('The Role of Cognitive Load in Second Language Acquisition', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(12)

    # --- Author and institution ---
    author_para = doc.add_paragraph('Dr. Elena Marchetti, Department of Applied Linguistics')
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_after = Pt(6)

    affiliation_para = doc.add_paragraph('University of Geneva, Geneva, Switzerland')
    affiliation_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affiliation_para.paragraph_format.space_after = Pt(18)

    # --- Abstract heading ---
    abstract_heading = doc.add_paragraph('Abstract')
    abstract_heading.runs[0].bold = True
    abstract_heading.paragraph_format.space_after = Pt(6)

    # --- Body paragraph 1 (1.15 line spacing — initial state) ---
    para1_text = (
        "Cognitive load theory, originally proposed by Sweller (1988), has profound implications "
        "for second language acquisition (SLA) research. When learners encounter novel linguistic "
        "structures in a target language, the working memory system must simultaneously process "
        "phonological, syntactic, and semantic information. This multidimensional processing demand "
        "creates what researchers term 'intrinsic cognitive load', which varies considerably depending "
        "on the learner's prior linguistic knowledge and the structural distance between the native "
        "language and the target language. Studies conducted across European university populations "
        "suggest that intrinsic load is significantly higher for learners whose native languages "
        "belong to different language families than the target language (Brysbaert & Duyck, 2010)."
    )
    para1 = doc.add_paragraph(para1_text)
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para1.paragraph_format.line_spacing = 1.15
    para1.paragraph_format.space_after = Pt(10)

    # --- Body paragraph 2 (1.15 line spacing — initial state) ---
    para2_text = (
        "In contrast, 'extraneous cognitive load' arises from the design of instructional materials "
        "and the pedagogical environment rather than from the inherent complexity of the content itself. "
        "Poorly structured grammar exercises, inconsistent orthographic conventions in teaching materials, "
        "and inadequate scaffolding in communicative tasks all contribute to extraneous load that "
        "competes for limited working memory resources. Schmidt's noticing hypothesis (1990) intersects "
        "meaningfully with cognitive load frameworks, suggesting that attentional resources diverted "
        "by extraneous factors reduce the likelihood that learners will consciously register the formal "
        "features of language input — a prerequisite for subsequent acquisition. Classroom-based "
        "interventions that minimize extraneous load through explicit metalinguistic instruction have "
        "demonstrated measurable gains in morphosyntactic accuracy among intermediate-level learners."
    )
    para2 = doc.add_paragraph(para2_text)
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para2.paragraph_format.line_spacing = 1.15
    para2.paragraph_format.space_after = Pt(10)

    # --- Body paragraph 3 (1.15 line spacing — initial state) ---
    para3_text = (
        "The third dimension of cognitive load, 'germane cognitive load', refers to the mental effort "
        "devoted to schema formation and automation — processes that are directly beneficial to long-term "
        "learning. In SLA contexts, germane load is cultivated through tasks that require learners to "
        "detect patterns, formulate hypotheses about target-language rules, and receive corrective feedback "
        "in a manner that promotes explicit rule representation. Doughty and Williams (1998) demonstrated "
        "that form-focused instruction embedded within meaning-oriented tasks optimally engages germane "
        "processing without overwhelming the learner's total cognitive capacity. More recent neuroimaging "
        "research corroborates these behavioural findings, indicating that activation in the left inferior "
        "frontal gyrus correlates with successful grammar learning under conditions of moderate germane load, "
        "suggesting a neurobiological basis for the optimal challenge hypothesis in language pedagogy."
    )
    para3 = doc.add_paragraph(para3_text)
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para3.paragraph_format.line_spacing = 1.15
    para3.paragraph_format.space_after = Pt(10)

    # --- References heading ---
    ref_heading = doc.add_paragraph('References')
    ref_heading.runs[0].bold = True
    ref_heading.paragraph_format.space_before = Pt(12)
    ref_heading.paragraph_format.space_after = Pt(6)

    # --- Reference entries ---
    refs = [
        'Brysbaert, M., & Duyck, W. (2010). Is it time to leave behind the Revised Hierarchical Model of bilingual language processing after fifteen years of service? Bilingualism: Language and Cognition, 13(3), 359–371.',
        'Doughty, C., & Williams, J. (1998). Pedagogical choices in focus on form. In C. Doughty & J. Williams (Eds.), Focus on Form in Classroom Second Language Acquisition (pp. 197–261). Cambridge University Press.',
        'Schmidt, R. (1990). The role of consciousness in second language learning. Applied Linguistics, 11(2), 129–158.',
        'Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257–285.',
    ]
    for ref_text in refs:
        ref_para = doc.add_paragraph(ref_text)
        ref_para.paragraph_format.first_line_indent = Inches(-0.25)
        ref_para.paragraph_format.left_indent = Inches(0.25)
        ref_para.paragraph_format.space_after = Pt(4)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
