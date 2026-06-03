"""
Reward Script: Set unique line spacing for each of the three body paragraphs
Task ID: osworld_writer_line_spacing_per_paragraph_007
Domain: libreoffice_writer
Scoring:
  Component 1: First body paragraph (para index 4) has line_spacing == 1.0 (SINGLE)  — 0.4 pts
  Component 2: Second body paragraph (para index 5) has line_spacing == 1.5 (1.5x)   — 0.3 pts
  Component 3: Third body paragraph (para index 6) has line_spacing == 2.0 (DOUBLE)  — 0.3 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.enum.text import WD_LINE_SPACING

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_007'

# The three body paragraphs in the academic paper are at indices 4, 5, 6:
#   Para 4 starts: "Cognitive load theory..."
#   Para 5 starts: "In contrast, 'extraneous cognitive load'..."
#   Para 6 starts: "The third dimension of cognitive load..."
# In the initial state they all have 1.15 spacing (MULTIPLE).
# The task requires: para4=1.0 (SINGLE), para5=1.5 (ONE_POINT_FIVE), para6=2.0 (DOUBLE).

BODY_PARA_INDICES = [4, 5, 6]

# Expected line spacing values
EXPECTED_SPACING = {
    4: 1.0,   # single
    5: 1.5,   # 1.5x
    6: 2.0,   # double
}

# Expected line spacing rules
EXPECTED_RULES = {
    4: WD_LINE_SPACING.SINGLE,          # 0
    5: WD_LINE_SPACING.ONE_POINT_FIVE,  # 1
    6: WD_LINE_SPACING.DOUBLE,          # 2
}

RULE_NAMES = {
    4: "SINGLE",
    5: "ONE_POINT_FIVE",
    6: "DOUBLE",
}

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Sanity check: document must have at least 7 paragraphs
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected at least 7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: First body paragraph (index 4) has line_spacing == 1.0 (SINGLE) (0.4 points)
    # This FAILS on initial (1.15/MULTIPLE) and PASSES on golden (1.0/SINGLE)
    try:
        para4 = doc.paragraphs[4]
        pf4 = para4.paragraph_format
        spacing4 = pf4.line_spacing
        rule4 = pf4.line_spacing_rule
        # Accept either: line_spacing==1.0 with SINGLE rule, OR just line_spacing==1.0
        # (some implementations may express it differently, so also accept numeric 240/240ths)
        if spacing4 is not None and abs(float(spacing4) - 1.0) < 0.05:
            print(f"PASS: Component 1 — Para 4 line_spacing={spacing4}, rule={rule4} (expected 1.0/SINGLE) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Para 4 line_spacing={spacing4}, rule={rule4} (expected 1.0/SINGLE)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Second body paragraph (index 5) has line_spacing == 1.5 (ONE_POINT_FIVE) (0.3 points)
    # This FAILS on initial (1.15/MULTIPLE) and PASSES on golden (1.5/ONE_POINT_FIVE)
    try:
        para5 = doc.paragraphs[5]
        pf5 = para5.paragraph_format
        spacing5 = pf5.line_spacing
        rule5 = pf5.line_spacing_rule
        if spacing5 is not None and abs(float(spacing5) - 1.5) < 0.05:
            print(f"PASS: Component 2 — Para 5 line_spacing={spacing5}, rule={rule5} (expected 1.5/ONE_POINT_FIVE) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Para 5 line_spacing={spacing5}, rule={rule5} (expected 1.5/ONE_POINT_FIVE)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Third body paragraph (index 6) has line_spacing == 2.0 (DOUBLE) (0.3 points)
    # This FAILS on initial (1.15/MULTIPLE) and PASSES on golden (2.0/DOUBLE)
    try:
        para6 = doc.paragraphs[6]
        pf6 = para6.paragraph_format
        spacing6 = pf6.line_spacing
        rule6 = pf6.line_spacing_rule
        if spacing6 is not None and abs(float(spacing6) - 2.0) < 0.05:
            print(f"PASS: Component 3 — Para 6 line_spacing={spacing6}, rule={rule6} (expected 2.0/DOUBLE) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Para 6 line_spacing={spacing6}, rule={rule6} (expected 2.0/DOUBLE)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
