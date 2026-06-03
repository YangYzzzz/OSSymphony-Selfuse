"""
Reward Script: Benefits Comparison Document with Multiple Tables
Task ID: writer_hr_064
Domain: libreoffice_writer
Scoring:
  Component 1: Health Plans table (5+1 rows x 8 cols) — 0.20 pts
  Component 2: Dental Plans table (3+1 rows x 6 cols) — 0.15 pts
  Component 3: Vision Plans table (2+1 rows x 5 cols) — 0.15 pts
  Component 4: Retirement Plans table (2+1 rows x 7 cols) — 0.15 pts
  Component 5: Table of Tables with 4 entries — 0.15 pts
  Component 6: Table captions present for all 4 tables — 0.10 pts
  Component 7: Header rows have bold formatting — 0.10 pts
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_064'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_tables = len(doc.tables)
    print(f"INFO: Document has {num_tables} tables and {len(doc.paragraphs)} paragraphs")

    # We need at least 5 tables: Table of Tables + 4 data tables
    if num_tables < 5:
        print(f"FAIL: Need at least 5 tables (Table of Tables + 4 data tables), found {num_tables}")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Identify the Table of Tables (ToT) and the 4 data tables.
    # ToT is the table that has entries like "Table 1", "Table 2", etc.
    # Data tables are the remaining ones.
    tot_index = None
    data_table_indices = []

    for t_idx, table in enumerate(doc.tables):
        # Check if this looks like a Table of Tables
        # It should have rows referencing "Table 1", "Table 2", etc.
        is_tot = False
        tot_refs = 0
        for row in table.rows:
            cell_texts = [cell.text.strip().lower() for cell in row.cells]
            joined = ' '.join(cell_texts)
            for ref in ['table 1', 'table 2', 'table 3', 'table 4']:
                if ref in joined:
                    tot_refs += 1
        if tot_refs >= 3:
            is_tot = True

        if is_tot and tot_index is None:
            tot_index = t_idx
        else:
            data_table_indices.append(t_idx)

    if tot_index is None:
        print("FAIL: Could not identify Table of Tables")
    else:
        print(f"INFO: Table of Tables found at index {tot_index}")

    # Map data tables by checking header content
    health_table = None
    dental_table = None
    vision_table = None
    retirement_table = None

    for t_idx in data_table_indices:
        table = doc.tables[t_idx]
        if len(table.rows) == 0:
            continue
        # Get header row text
        header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
        header_joined = ' '.join(header_texts)

        if 'plan name' in header_joined and ('premium' in header_joined or 'deductible' in header_joined):
            # Could be health, dental, or vision
            num_cols = len(table.columns)
            num_rows = len(table.rows)
            if num_cols >= 7 and num_rows >= 5:
                health_table = table
            elif 'preventive' in header_joined or 'major procedures' in header_joined:
                dental_table = table
            elif 'eye exam' in header_joined or 'frames' in header_joined or 'contact' in header_joined:
                vision_table = table
            elif num_cols >= 5 and num_cols <= 6 and num_rows >= 3:
                dental_table = table
            elif num_cols <= 5 and num_rows <= 3:
                vision_table = table
        elif 'feature' in header_joined or 'traditional' in header_joined or '401' in header_joined or 'roth' in header_joined:
            retirement_table = table
        elif 'eye exam' in header_joined or 'frames' in header_joined or 'contact' in header_joined:
            vision_table = table
        elif 'preventive' in header_joined or 'basic procedures' in header_joined:
            dental_table = table

    # Component 1: Health Plans table (5+1 rows x 8 cols) — 0.20 pts
    try:
        if health_table is not None:
            h_rows = len(health_table.rows)
            h_cols = len(health_table.columns)
            # Expect at least 6 rows (1 header + 5 data) and 8 columns
            if h_rows >= 6 and h_cols >= 8:
                print(f"PASS: Component 1 — Health Plans table found ({h_rows} rows x {h_cols} cols) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 1 — Health Plans table wrong dimensions: {h_rows} rows x {h_cols} cols (need >=6 rows, >=8 cols)")
        else:
            print("FAIL: Component 1 — Health Plans table not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Dental Plans table (3+1 rows x 6 cols) — 0.15 pts
    try:
        if dental_table is not None:
            d_rows = len(dental_table.rows)
            d_cols = len(dental_table.columns)
            # Expect at least 4 rows (1 header + 3 data) and 6 columns
            if d_rows >= 4 and d_cols >= 6:
                print(f"PASS: Component 2 — Dental Plans table found ({d_rows} rows x {d_cols} cols) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — Dental Plans table wrong dimensions: {d_rows} rows x {d_cols} cols (need >=4 rows, >=6 cols)")
        else:
            print("FAIL: Component 2 — Dental Plans table not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Vision Plans table (2+1 rows x 5 cols) — 0.15 pts
    try:
        if vision_table is not None:
            v_rows = len(vision_table.rows)
            v_cols = len(vision_table.columns)
            # Expect at least 3 rows (1 header + 2 data) and 5 columns
            if v_rows >= 3 and v_cols >= 5:
                print(f"PASS: Component 3 — Vision Plans table found ({v_rows} rows x {v_cols} cols) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 — Vision Plans table wrong dimensions: {v_rows} rows x {v_cols} cols (need >=3 rows, >=5 cols)")
        else:
            print("FAIL: Component 3 — Vision Plans table not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Retirement Plans table (2+1 rows x 7 cols) — 0.15 pts
    try:
        if retirement_table is not None:
            r_rows = len(retirement_table.rows)
            r_cols = len(retirement_table.columns)
            # Expect at least 3 rows (1 header + 2 data) and 7 columns
            if r_rows >= 3 and r_cols >= 7:
                print(f"PASS: Component 4 — Retirement Plans table found ({r_rows} rows x {r_cols} cols) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — Retirement Plans table wrong dimensions: {r_rows} rows x {r_cols} cols (need >=3 rows, >=7 cols)")
        else:
            print("FAIL: Component 4 — Retirement Plans table not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Table of Tables with 4 entries — 0.15 pts
    try:
        if tot_index is not None:
            tot_table = doc.tables[tot_index]
            # Count data rows (excluding header)
            data_rows = len(tot_table.rows) - 1  # minus header row
            if data_rows >= 4:
                # Verify entries reference the 4 tables
                refs_found = set()
                for row in tot_table.rows[1:]:
                    cell_text = ' '.join(cell.text.strip().lower() for cell in row.cells)
                    for ref_num in ['1', '2', '3', '4']:
                        if f'table {ref_num}' in cell_text or f'table{ref_num}' in cell_text:
                            refs_found.add(ref_num)
                if len(refs_found) >= 4:
                    print(f"PASS: Component 5 — Table of Tables has {data_rows} entries referencing tables {refs_found} (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 5 — Table of Tables has entries but only references tables: {refs_found}")
            else:
                print(f"FAIL: Component 5 — Table of Tables has only {data_rows} data rows (need >=4)")
        else:
            print("FAIL: Component 5 — Table of Tables not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Table captions present for all 4 tables — 0.10 pts
    # Captions should appear as paragraphs containing "Table N:" or "Table N." before each table
    try:
        para_texts = [p.text.strip().lower() for p in doc.paragraphs]
        captions_found = 0
        caption_keywords = [
            'table 1',  # Health
            'table 2',  # Dental
            'table 3',  # Vision
            'table 4',  # Retirement
        ]
        for kw in caption_keywords:
            # Check if any paragraph contains this caption reference
            found = False
            for pt in para_texts:
                if kw in pt and ('plan' in pt or 'comparison' in pt or 'insurance' in pt or 'dental' in pt or 'vision' in pt or 'retirement' in pt or 'saving' in pt):
                    found = True
                    break
            if found:
                captions_found += 1

        if captions_found >= 4:
            print(f"PASS: Component 6 — All 4 table captions found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — Only {captions_found}/4 table captions found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Header rows have bold formatting in data tables — 0.10 pts
    # Check that at least 3 of the 4 data tables have bold headers
    try:
        tables_with_bold_headers = 0
        for label, tbl in [('Health', health_table), ('Dental', dental_table),
                           ('Vision', vision_table), ('Retirement', retirement_table)]:
            if tbl is None:
                continue
            header_row = tbl.rows[0]
            bold_count = 0
            total_cells = 0
            for cell in header_row.cells:
                total_cells += 1
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip() and run.font.bold:
                            bold_count += 1
                            break
            if bold_count >= total_cells * 0.5:  # At least half the header cells are bold
                tables_with_bold_headers += 1

        if tables_with_bold_headers >= 3:
            print(f"PASS: Component 7 — {tables_with_bold_headers}/4 data tables have bold headers (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 — Only {tables_with_bold_headers}/4 data tables have bold headers (need >=3)")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
