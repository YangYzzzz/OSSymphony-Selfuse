"""
Initial Setup: Benefits Comparison Document with intro only
Task ID: writer_hr_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_064'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('Benefits Comparison Guide 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technology Solutions, Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.italic = True

    # Date line
    date_line = doc.add_paragraph()
    date_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_line.add_run('Effective January 1, 2026 — December 31, 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()  # spacer

    # Introduction heading
    doc.add_heading('Introduction', level=1)

    # Introduction paragraphs
    intro1 = doc.add_paragraph(
        'Welcome to the 2026 Benefits Comparison Guide for Meridian Technology Solutions. '
        'This document provides a comprehensive overview of all benefit plans available to '
        'full-time employees and their eligible dependents. Our benefits package has been '
        'designed to support your health, financial security, and overall well-being.'
    )
    intro1.paragraph_format.space_after = Pt(8)

    intro2 = doc.add_paragraph(
        'During Open Enrollment (October 15 – November 15, 2025), you may review and select '
        'the plans that best meet your needs. Changes made during Open Enrollment take effect '
        'on January 1, 2026. If you do not make changes, your current elections will roll over '
        'into the new plan year, except where noted.'
    )
    intro2.paragraph_format.space_after = Pt(8)

    intro3 = doc.add_paragraph(
        'Please review each section carefully. For questions about specific plan details, '
        'contact the HR Benefits Team at benefits@meridiantech.com or extension 4200. '
        'You may also schedule a one-on-one consultation during Open Enrollment weeks.'
    )
    intro3.paragraph_format.space_after = Pt(8)

    # Overview heading
    doc.add_heading('Benefits Overview', level=1)

    overview = doc.add_paragraph(
        'Meridian Technology Solutions offers four categories of benefits to eligible employees: '
        'Health Insurance, Dental Coverage, Vision Care, and Retirement Savings. Each category '
        'includes multiple plan options to accommodate different coverage needs and budget '
        'preferences. The following sections contain detailed comparison tables for each '
        'benefit category.'
    )
    overview.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
