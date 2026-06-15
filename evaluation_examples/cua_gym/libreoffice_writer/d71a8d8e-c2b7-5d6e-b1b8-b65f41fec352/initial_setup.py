"""
Initial Setup: Company policy document with dense paragraph (8 sentences, no blank lines between them)
Task ID: osworld_writer_blank_line_insertion_007
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_blank_line_insertion_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title heading ---
    title = doc.add_heading('Remote Work Policy', level=1)

    # --- Subtitle / metadata ---
    meta = doc.add_paragraph('Effective Date: January 1, 2025 | Department: Human Resources | Version: 3.2')
    meta.paragraph_format.space_after = Pt(12)

    # --- Section heading ---
    doc.add_heading('Policy Overview', level=2)

    # --- Dense body paragraph with 8 sentences separated only by spaces ---
    # NO blank lines between sentences — that is what the agent must add
    sentences = (
        "This Remote Work Policy establishes the guidelines and expectations for employees who work outside of the primary office location. "
        "All eligible employees must submit a formal remote work agreement to their direct manager and receive written approval prior to commencing any remote arrangement. "
        "Employees are expected to maintain their standard working hours and remain reachable via company-approved communication channels throughout the business day. "
        "The company reserves the right to rescind remote work privileges at any time if performance standards are not consistently met or if business needs require on-site presence. "
        "Each employee working remotely is responsible for ensuring a safe, ergonomically appropriate, and distraction-free workspace at their designated remote location. "
        "All confidential company data and client information must be handled in strict accordance with the Information Security Policy, which includes the use of encrypted connections and VPN access. "
        "IT support for remote employees is available during standard business hours, and any hardware or software issues must be reported through the official helpdesk ticketing system. "
        "Compliance with this policy is mandatory for all staff in roles approved for remote work, and violations may result in disciplinary action up to and including termination of the remote work arrangement."
    )

    body_para = doc.add_paragraph(sentences)
    body_para.paragraph_format.space_before = Pt(6)
    body_para.paragraph_format.space_after = Pt(6)

    # --- Footer section ---
    doc.add_heading('Acknowledgment', level=2)
    ack = doc.add_paragraph(
        'By continuing to work under this arrangement, employees confirm they have read, understood, and agree to comply with the terms outlined in this Remote Work Policy.'
    )
    ack.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
