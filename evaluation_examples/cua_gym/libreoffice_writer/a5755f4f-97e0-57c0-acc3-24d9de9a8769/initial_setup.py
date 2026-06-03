"""
Initial Setup: Sales rankings table (unsorted) in sales_rankings.docx
Task ID: writer_tbl_025
Domain: libreoffice_writer

Creates a .docx file with a 6-row x 3-column table containing sales data
in its original (unsorted) order. The agent task is to sort the data rows
by 'Total Sales' descending.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'writer_tbl_025'
FILENAME = 'sales_rankings.docx'
OUTPUT = f'{DESKTOP}/{FILENAME}'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Add a title paragraph for context
    title = doc.add_paragraph('Sales Representative Rankings')
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacing

    # Create the table: 6 rows (1 header + 5 data) x 3 columns
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    # Row 0 — header
    headers = ['Sales Rep', 'Region', 'Total Sales']
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True

    # Data rows — UNSORTED (original order as provided in context)
    data_rows = [
        ('Alice',  'West',    '45000'),
        ('Bob',    'East',    '62000'),
        ('Carol',  'North',   '38000'),
        ('Dave',   'South',   '71000'),
        ('Eve',    'Central', '55000'),
    ]

    for row_idx, (rep, region, sales) in enumerate(data_rows, start=1):
        table.cell(row_idx, 0).text = rep
        table.cell(row_idx, 1).text = region
        table.cell(row_idx, 2).text = sales

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
