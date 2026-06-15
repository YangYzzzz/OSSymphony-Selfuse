"""
Reward Script: Sort table data by 'Total Sales' column in descending order
Task ID: writer_tbl_025
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4 pts): First data row is Dave/South/71000 (highest Total Sales value)
  Component 2 (0.6 pts): All 5 data rows are in the correct descending order by Total Sales
Total: 1.0
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_025'

# Expected sort order (descending by Total Sales, header excluded)
EXPECTED_SORTED_ROWS = [
    ('Dave', 'South', '71000'),
    ('Bob', 'East', '62000'),
    ('Eve', 'Central', '55000'),
    ('Alice', 'West', '45000'),
    ('Carol', 'North', '38000'),
]

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Sort the table data by 'Total Sales' column in descending order.
    Initial state: rows in arbitrary order (Alice, Bob, Carol, Dave, Eve)
    Golden state: rows sorted descending by Total Sales (Dave 71000, Bob 62000, Eve 55000, Alice 45000, Carol 38000)
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: file must have at least one table with 6 rows
    if not doc.tables:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) < 6:
        print(f"FAIL: Expected 6 rows (1 header + 5 data), found {len(table.rows)}")
        print("REWARD: 0.0")
        return 0.0

    # Extract data rows (skip header row 0)
    data_rows = []
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        data_rows.append(tuple(cells))

    print(f"Data rows found: {data_rows}")

    # Component 1: First data row is Dave/South/71000 (the highest Total Sales entry) (0.4 points)
    # This FAILS on initial_env (first data row = Alice/West/45000) and PASSES on golden_env
    try:
        first_row = data_rows[0] if data_rows else ()
        expected_first = ('Dave', 'South', '71000')
        if first_row == expected_first:
            print(f"PASS: Component 1 — First data row is Dave/South/71000 (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected first data row {expected_first}, found {first_row}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 5 data rows are in the correct descending order by Total Sales (0.6 points)
    # This FAILS on initial_env (original unsorted order) and PASSES on golden_env
    try:
        if len(data_rows) == 5 and data_rows == EXPECTED_SORTED_ROWS:
            print(f"PASS: Component 2 — All 5 data rows in correct descending order by Total Sales (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 2 — Expected sorted order {EXPECTED_SORTED_ROWS}, found {data_rows}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = f'{WORKDIR}/Desktop/sales_rankings.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
