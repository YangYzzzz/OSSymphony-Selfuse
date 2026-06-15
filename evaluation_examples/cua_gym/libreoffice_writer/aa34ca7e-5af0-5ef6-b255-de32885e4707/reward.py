"""
Reward Script: Invoice template creation in LibreOffice Writer
Task ID: writer_wf_019
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Company name 'Apex Consulting Group' in 20pt bold
  Component 2 (0.20): Invoice info fields (Invoice #, Date, Due Date, Client Name, Client Address)
  Component 3 (0.25): Table with 4 columns (Description, Hours, Rate, Amount) and 4+ data rows
  Component 4 (0.20): Subtotal, Tax (10%), and Total rows
  Component 5 (0.15): Payment terms section at bottom
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_019'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify invoice template creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_para_texts = [p.text.strip() for p in doc.paragraphs]
    non_empty_paras = [t for t in all_para_texts if t]

    # Component 1: Company name 'Apex Consulting Group' in 20pt bold (0.20 points)
    try:
        company_found = False
        for para in doc.paragraphs:
            if 'apex consulting group' in para.text.lower():
                for run in para.runs:
                    if 'apex consulting group' in run.text.lower():
                        is_bold = run.font.bold is True
                        size_pt = run.font.size.pt if run.font.size else None
                        # Allow some tolerance on size (19-21pt)
                        size_ok = size_pt is not None and 19.0 <= size_pt <= 21.0
                        if is_bold and size_ok:
                            company_found = True
                            print(f"PASS: Component 1 — 'Apex Consulting Group' found, bold={is_bold}, size={size_pt}pt (0.20 pts)")
                            total_score += 0.20
                            break
                        else:
                            print(f"FAIL: Component 1 — 'Apex Consulting Group' found but bold={is_bold}, size={size_pt}pt (expected bold=True, ~20pt)")
                if company_found:
                    break
        if not company_found:
            print("FAIL: Component 1 — 'Apex Consulting Group' not found in 20pt bold")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Invoice info fields present (0.20 points)
    # Check for: Invoice #, Date, Due Date, Client Name, Client Address
    try:
        full_text = '\n'.join(all_para_texts).lower()
        info_fields = {
            'invoice #': False,
            'date': False,
            'due date': False,
            'client name': False,
            'client address': False,
        }
        for para in doc.paragraphs:
            text_lower = para.text.lower().strip()
            if 'invoice' in text_lower and '#' in text_lower:
                info_fields['invoice #'] = True
            if text_lower.startswith('date') or ': ' in text_lower and text_lower.split(':')[0].strip() == 'date':
                info_fields['date'] = True
            if 'due date' in text_lower:
                info_fields['due date'] = True
            if 'client name' in text_lower:
                info_fields['client name'] = True
            if 'client address' in text_lower:
                info_fields['client address'] = True

        # Also check if 'date' appears in general context (some might format differently)
        if not info_fields['date'] and 'date' in full_text:
            info_fields['date'] = True

        fields_found = sum(1 for v in info_fields.values() if v)
        # Need at least 4 of 5 for full credit, proportional otherwise
        if fields_found >= 4:
            print(f"PASS: Component 2 — {fields_found}/5 info fields found: {info_fields} (0.20 pts)")
            total_score += 0.20
        elif fields_found >= 2:
            partial = round(0.20 * fields_found / 5, 2)
            print(f"PARTIAL: Component 2 — {fields_found}/5 info fields found: {info_fields} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {fields_found}/5 info fields found: {info_fields}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table with 4 columns and 4+ data rows (0.25 points)
    try:
        table_score = 0.0
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_cols = len(table.columns)
            num_rows = len(table.rows)

            # Check column headers
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            expected_headers = ['description', 'hours', 'rate', 'amount']
            headers_match = sum(1 for eh in expected_headers if any(eh in hc for hc in header_cells))

            # Check data rows (excluding header)
            data_rows = num_rows - 1  # subtract header row

            if num_cols >= 4 and headers_match >= 3:
                table_score += 0.10
                print(f"  Table structure: {num_cols} cols, headers match {headers_match}/4")

            if data_rows >= 4:
                table_score += 0.10
                print(f"  Data rows: {data_rows} (need >= 4)")
            elif data_rows >= 2:
                table_score += 0.05
                print(f"  Data rows: {data_rows} (partial, need >= 4)")

            # Check that data rows have meaningful content (not empty)
            non_empty_data_rows = 0
            for ri in range(1, num_rows):
                row_text = ''.join(cell.text.strip() for cell in table.rows[ri].cells)
                if row_text:
                    non_empty_data_rows += 1
            if non_empty_data_rows >= 4:
                table_score += 0.05
                print(f"  Non-empty data rows: {non_empty_data_rows}")

            if table_score > 0:
                print(f"PASS: Component 3 — Table found with {num_cols} cols, {data_rows} data rows ({table_score} pts)")
            else:
                print(f"FAIL: Component 3 — Table found but cols={num_cols}, data_rows={data_rows}, headers_match={headers_match}")
            total_score += table_score
        else:
            print("FAIL: Component 3 — No table found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Subtotal, Tax (10%), and Total rows (0.20 points)
    try:
        subtotal_found = False
        tax_found = False
        total_found = False

        for para in doc.paragraphs:
            text_lower = para.text.lower().strip()
            if 'subtotal' in text_lower:
                subtotal_found = True
            if 'tax' in text_lower and '10%' in text_lower:
                tax_found = True
            # "total" but not "subtotal" — check that line starts or contains standalone "total"
            if text_lower.startswith('total') or (': ' in text_lower and 'total' in text_lower.split(':')[0].lower() and 'subtotal' not in text_lower):
                total_found = True

        summary_count = sum([subtotal_found, tax_found, total_found])
        if summary_count == 3:
            print(f"PASS: Component 4 — Subtotal={subtotal_found}, Tax(10%)={tax_found}, Total={total_found} (0.20 pts)")
            total_score += 0.20
        elif summary_count >= 1:
            partial = round(0.20 * summary_count / 3, 2)
            print(f"PARTIAL: Component 4 — Subtotal={subtotal_found}, Tax(10%)={tax_found}, Total={total_found} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No subtotal/tax/total rows found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Payment terms section at bottom (0.15 points)
    try:
        payment_found = False
        # Check the last few non-empty paragraphs for payment terms
        for para in doc.paragraphs[-5:]:
            text_lower = para.text.lower().strip()
            if 'payment' in text_lower and ('terms' in text_lower or 'due' in text_lower or 'days' in text_lower):
                payment_found = True
                break

        if payment_found:
            print(f"PASS: Component 5 — Payment terms section found at bottom (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — No payment terms found in last paragraphs")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
