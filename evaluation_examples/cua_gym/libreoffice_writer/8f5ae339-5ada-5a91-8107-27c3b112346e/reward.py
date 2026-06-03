"""
Reward Script: Create a professional CV/resume for Alex Chen in LibreOffice Writer
Task ID: writer_wf_009
Domain: libreoffice_writer
Scoring:
  Component 1: Name "Alex Chen" large and bold at top (0.15)
  Component 2: Contact information with email, phone, LinkedIn (0.10)
  Component 3: Professional Summary section with content (0.15)
  Component 4: Work Experience section with 2 positions (0.20)
  Component 5: Bullet points in work experience (0.10)
  Component 6: Education section present (0.10)
  Component 7: Skills table — 2 columns, 4 rows, correct categories (0.20)
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_009'


def persist_app_state(domain: str):
    """Try to save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    tables = doc.tables

    # Precondition: document must have content (blank doc = 0.0)
    if len(paragraphs) == 0:
        print("FAIL: Document is empty (0 paragraphs)")
        print("REWARD: 0.0")
        return 0.0

    # Collect full text for section detection
    full_texts = [p.text.strip() for p in paragraphs]
    full_texts_lower = [t.lower() for t in full_texts]

    # Component 1: Name "Alex Chen" large and bold at top (0.15 points)
    try:
        first_para = paragraphs[0]
        first_text = first_para.text.strip()
        name_found = 'alex chen' in first_text.lower()
        # Check if bold
        has_bold = any(r.bold for r in first_para.runs if r.text.strip())
        # Check if large font (>= 16pt considered large)
        has_large = False
        for r in first_para.runs:
            if r.font.size and r.font.size.pt >= 16:
                has_large = True
                break

        if name_found and has_bold and has_large:
            sizes = [r.font.size.pt for r in first_para.runs if r.font.size]
            print(f"PASS: Component 1 — 'Alex Chen' found bold at top, font sizes: {sizes} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — name_found={name_found}, has_bold={has_bold}, has_large={has_large}, text='{first_text[:60]}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Contact information (email, phone, LinkedIn) (0.10 points)
    try:
        contact_found = False
        email_ok = False
        phone_ok = False
        linkedin_ok = False
        # Search first 5 paragraphs for contact info
        for p in paragraphs[:5]:
            text = p.text.lower()
            if '@' in text or 'email' in text:
                email_ok = True
            if any(c.isdigit() for c in text) and ('(' in text or '-' in text or '+' in text):
                phone_ok = True
            if 'linkedin' in text:
                linkedin_ok = True

        if email_ok and phone_ok and linkedin_ok:
            print(f"PASS: Component 2 — Contact info found: email={email_ok}, phone={phone_ok}, linkedin={linkedin_ok} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — email={email_ok}, phone={phone_ok}, linkedin={linkedin_ok}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Professional Summary section with content (0.15 points)
    try:
        summary_idx = None
        for i, t in enumerate(full_texts_lower):
            if 'professional summary' in t or 'summary' == t or 'profile' == t:
                summary_idx = i
                break

        if summary_idx is not None:
            # Check there is content after the heading
            has_content = False
            for j in range(summary_idx + 1, min(summary_idx + 4, len(paragraphs))):
                text = paragraphs[j].text.strip()
                # Stop if we hit the next section heading
                if text.upper() == text and len(text) > 3 and text.isalpha():
                    break
                if len(text) > 30:  # meaningful content
                    has_content = True
                    break

            if has_content:
                print(f"PASS: Component 3 — Professional Summary section found with content (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 — Summary section heading found but no content after it")
        else:
            print(f"FAIL: Component 3 — No 'Professional Summary' section heading found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Work Experience section with 2 positions (0.20 points)
    try:
        work_idx = None
        for i, t in enumerate(full_texts_lower):
            if 'work experience' in t or 'experience' == t:
                work_idx = i
                break

        if work_idx is not None:
            # Count bold title paragraphs after work experience heading (job titles are bold)
            # Also look for date patterns as indicators of positions
            position_count = 0
            # Find the next major section heading after work experience
            next_section_idx = len(paragraphs)
            section_keywords = ['education', 'skills', 'certifications', 'projects', 'references']
            for i in range(work_idx + 1, len(paragraphs)):
                t = full_texts_lower[i]
                if any(kw == t or kw in t for kw in section_keywords):
                    # Check if it looks like a heading (bold, short)
                    if len(paragraphs[i].text.strip()) < 30:
                        has_bold_run = any(r.bold for r in paragraphs[i].runs if r.text.strip())
                        if has_bold_run:
                            next_section_idx = i
                            break

            # Count positions by finding bold non-bullet paragraphs (job titles)
            for i in range(work_idx + 1, next_section_idx):
                p = paragraphs[i]
                style_name = p.style.name if p.style else ''
                if 'Bullet' in style_name or 'List' in style_name:
                    continue
                text = p.text.strip()
                if not text:
                    continue
                has_bold_run = any(r.bold for r in p.runs if r.text.strip())
                # Position titles are bold and not too long
                if has_bold_run and len(text) < 60:
                    position_count += 1

            if position_count >= 2:
                print(f"PASS: Component 4 — Work Experience with {position_count} positions found (0.20 pts)")
                total_score += 0.20
            elif position_count == 1:
                print(f"PARTIAL: Component 4 — Only {position_count} position found (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Work Experience section found but {position_count} positions detected")
        else:
            print(f"FAIL: Component 4 — No 'Work Experience' section heading found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Bullet points in work experience (0.10 points)
    try:
        bullet_count = 0
        for p in paragraphs:
            style_name = p.style.name if p.style else ''
            if 'Bullet' in style_name or 'List Bullet' in style_name:
                if len(p.text.strip()) > 10:
                    bullet_count += 1

        if bullet_count >= 4:
            print(f"PASS: Component 5 — {bullet_count} bullet points found (0.10 pts)")
            total_score += 0.10
        elif bullet_count >= 2:
            print(f"PARTIAL: Component 5 — Only {bullet_count} bullet points found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — Only {bullet_count} bullet points found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Education section (0.10 points)
    try:
        education_idx = None
        for i, t in enumerate(full_texts_lower):
            if 'education' in t:
                education_idx = i
                break

        if education_idx is not None:
            # Check for degree information after heading
            has_degree = False
            for j in range(education_idx + 1, min(education_idx + 5, len(paragraphs))):
                text = paragraphs[j].text.lower()
                if any(kw in text for kw in ['bachelor', 'master', 'phd', 'degree', 'b.s.', 'b.a.', 'm.s.', 'science', 'arts']):
                    has_degree = True
                    break

            if has_degree:
                print(f"PASS: Component 6 — Education section with degree found (0.10 pts)")
                total_score += 0.10
            else:
                print(f"PARTIAL: Component 6 — Education section found but no degree info (0.05 pts)")
                total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No 'Education' section found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Skills table — 2 columns, 4 rows, correct categories (0.20 points)
    try:
        skills_idx = None
        for i, t in enumerate(full_texts_lower):
            if 'skills' == t or 'skills' in t:
                skills_idx = i
                break

        if skills_idx is None:
            print(f"FAIL: Component 7 — No 'Skills' section found")
        elif len(tables) == 0:
            print(f"FAIL: Component 7 — Skills section found but no tables in document")
        else:
            # Find the skills table (likely the last or only table)
            skills_table = tables[-1]  # assume last table is skills
            num_rows = len(skills_table.rows)
            num_cols = len(skills_table.columns)

            score_7 = 0.0

            # Check 2-column structure
            if num_cols == 2:
                score_7 += 0.05
                print(f"  Skills table: 2 columns confirmed")
            else:
                print(f"  Skills table: expected 2 cols, found {num_cols}")

            # Check 4 rows
            if num_rows >= 4:
                score_7 += 0.05
                print(f"  Skills table: {num_rows} rows (>= 4)")
            else:
                print(f"  Skills table: expected >= 4 rows, found {num_rows}")

            # Check category names in first column
            expected_categories = {'programming', 'frameworks', 'tools', 'soft skills'}
            actual_categories = set()
            for row in skills_table.rows:
                cat = row.cells[0].text.strip().lower()
                actual_categories.add(cat)

            matched = expected_categories & actual_categories
            if len(matched) >= 4:
                score_7 += 0.10
                print(f"  Skills table: all 4 categories found: {matched}")
            elif len(matched) >= 2:
                score_7 += 0.05
                print(f"  Skills table: {len(matched)}/4 categories found: {matched}")
            else:
                print(f"  Skills table: only {len(matched)} categories matched. Found: {actual_categories}")

            if score_7 > 0:
                print(f"PASS: Component 7 — Skills table verified ({score_7:.2f} pts)")
                total_score += score_7
            else:
                print(f"FAIL: Component 7 — Skills table structure incorrect")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
