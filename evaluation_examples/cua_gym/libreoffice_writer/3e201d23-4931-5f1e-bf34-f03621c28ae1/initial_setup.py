"""
Initial Setup: Insert a Table of Figures at the end of the document
Task ID: writer_rd_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_placeholder_image(filepath, label, width=640, height=400):
    """Create a simple placeholder diagram image with a label."""
    img = Image.new('RGB', (width, height), color=(230, 240, 250))
    draw = ImageDraw.Draw(img)
    # Draw border
    draw.rectangle([5, 5, width - 6, height - 6], outline=(100, 130, 170), width=2)
    # Draw some diagram-like shapes
    draw.rectangle([50, 50, 200, 150], outline=(70, 100, 150), width=2, fill=(200, 220, 245))
    draw.rectangle([300, 80, 450, 160], outline=(70, 100, 150), width=2, fill=(200, 220, 245))
    draw.rectangle([150, 220, 350, 320], outline=(70, 100, 150), width=2, fill=(200, 220, 245))
    draw.line([200, 100, 300, 120], fill=(70, 100, 150), width=2)
    draw.line([250, 160, 250, 220], fill=(70, 100, 150), width=2)
    # Label text
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
    except:
        font = ImageFont.load_default()
    draw.text((width // 2 - 80, height - 50), label, fill=(60, 60, 60), font=font)
    img.save(filepath)


def add_caption_paragraph(doc, caption_text):
    """Add a properly formatted figure caption using the 'Figure' SEQ field.
    This creates an SEQ field code that LibreOffice recognizes for Table of Figures."""
    para = doc.add_paragraph()
    para.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else doc.styles['Normal']
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add "Figure " text
    run_label = para.add_run("Figure ")
    run_label.font.size = Pt(10)
    run_label.italic = True

    # Add SEQ field for auto-numbering: this is what LibreOffice uses to build Table of Figures
    # Field begin
    r_begin = para.add_run()
    fldChar_begin = r_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r_begin._element.append(fldChar_begin)

    # Field instruction
    r_instr = para.add_run()
    instrText = r_instr._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instrText.text = ' SEQ Figure \\* ARABIC '
    r_instr._element.append(instrText)

    # Field separate
    r_sep = para.add_run()
    fldChar_sep = r_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r_sep._element.append(fldChar_sep)

    # Cached field result (the number)
    r_result = para.add_run(caption_text.split(":")[0].replace("Figure ", "").strip())
    r_result.font.size = Pt(10)
    r_result.italic = True

    # Field end
    r_end = para.add_run()
    fldChar_end = r_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r_end._element.append(fldChar_end)

    # Add the rest of the caption text (": Description")
    colon_and_desc = ":" + caption_text.split(":", 1)[1] if ":" in caption_text else ""
    run_desc = para.add_run(colon_and_desc)
    run_desc.font.size = Pt(10)
    run_desc.italic = True

    return para


def create_initial():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('Enterprise Cloud Infrastructure', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Technical Architecture Report')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(70, 100, 150)

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info.add_run('Prepared by: Infrastructure Engineering Team\nVersion 2.4 — March 2025\nClassification: Internal Use Only')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # --- Table of Contents placeholder (text only, not auto-generated) ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary .............. 3',
        '2. Network Architecture ........... 4',
        '3. Server Infrastructure .......... 5',
        '4. Database Design ................ 7',
        '5. API Gateway .................... 9',
        '6. Load Balancing Strategy ........ 11',
        '7. Monitoring & Observability ..... 13',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document provides a comprehensive overview of the enterprise cloud '
        'infrastructure deployed across three geographic regions. The architecture '
        'supports over 2.3 million daily active users and processes approximately '
        '450,000 API requests per minute during peak hours.'
    )
    doc.add_paragraph(
        'The infrastructure was redesigned in Q4 2024 to address scalability concerns '
        'identified during the annual capacity review. Key improvements include a '
        'transition to a service mesh topology, implementation of multi-region database '
        'replication, and deployment of an intelligent load balancing layer that reduces '
        'p99 latency by 34% compared to the previous architecture.'
    )

    # --- Section 2: Network Architecture ---
    doc.add_heading('2. Network Architecture', level=1)
    doc.add_paragraph(
        'The network layer implements a hub-and-spoke model with dedicated transit '
        'gateways connecting three regional virtual private clouds (VPCs). Each region '
        'maintains isolated subnets for public-facing services, internal microservices, '
        'and database clusters. Inter-region traffic is encrypted using TLS 1.3 with '
        'certificate pinning enforced at the service mesh layer.'
    )
    doc.add_paragraph(
        'Network segmentation follows a zero-trust model where each service must '
        'authenticate and authorize every request, regardless of source network. '
        'Security groups and network ACLs are managed through infrastructure-as-code '
        'templates, reviewed quarterly by the security operations team.'
    )

    # Figure 1
    img1_path = f'{WORKDIR}/fig_network_topology.png'
    create_placeholder_image(img1_path, 'Network Topology')
    doc.add_picture(img1_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_caption_paragraph(doc, 'Figure 1: Network Topology')

    doc.add_paragraph(
        'As shown in Figure 1, the three regional hubs connect through dedicated '
        '10 Gbps transit links with automatic failover to public internet paths. '
        'The topology ensures that no single point of failure can disconnect an '
        'entire region from the global service mesh.'
    )

    # --- Section 3: Server Infrastructure ---
    doc.add_heading('3. Server Infrastructure', level=1)
    doc.add_paragraph(
        'The compute layer runs on a Kubernetes-based container orchestration platform '
        'spanning 847 nodes across all regions. Each node pool is configured with '
        'auto-scaling policies that respond to CPU utilization, memory pressure, and '
        'custom application metrics published to the monitoring platform.'
    )
    doc.add_paragraph(
        'Production workloads are distributed across three tiers: front-end web servers '
        'handling HTTP/2 traffic, API gateway nodes processing business logic, and '
        'background workers managing asynchronous jobs such as email delivery, report '
        'generation, and data pipeline execution.'
    )

    # Figure 2
    img2_path = f'{WORKDIR}/fig_server_arch.png'
    create_placeholder_image(img2_path, 'Server Architecture', 640, 450)
    doc.add_picture(img2_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_caption_paragraph(doc, 'Figure 2: Server Architecture')

    doc.add_paragraph(
        'Figure 2 illustrates the multi-tier server architecture with clear separation '
        'between stateless application servers and stateful data services. The container '
        'runtime uses gVisor sandboxing for enhanced security isolation.'
    )

    # --- Section 4: Database Design ---
    doc.add_heading('4. Database Design', level=1)
    doc.add_paragraph(
        'The persistence layer employs a polyglot database strategy. Transactional data '
        'resides in a sharded PostgreSQL cluster with synchronous replication within each '
        'region and asynchronous cross-region replication. The primary cluster handles '
        'approximately 12,000 transactions per second with an average query latency of '
        '4.2 milliseconds.'
    )
    doc.add_paragraph(
        'Session state and caching leverage a Redis cluster with 256 GB of allocated '
        'memory per region. Search functionality is powered by an Elasticsearch cluster '
        'indexing 3.8 billion documents with sub-second query performance for 95% of '
        'search patterns.'
    )

    # Figure 3
    img3_path = f'{WORKDIR}/fig_database_schema.png'
    create_placeholder_image(img3_path, 'Database Schema', 640, 380)
    doc.add_picture(img3_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_caption_paragraph(doc, 'Figure 3: Database Schema')

    doc.add_paragraph(
        'The entity-relationship diagram in Figure 3 shows the core data model with '
        'twelve primary tables and their foreign key relationships. Denormalized read '
        'replicas are maintained for reporting and analytics workloads.'
    )

    doc.add_page_break()

    # --- Section 5: API Gateway ---
    doc.add_heading('5. API Gateway', level=1)
    doc.add_paragraph(
        'All external traffic enters through a unified API gateway that handles '
        'authentication, rate limiting, request validation, and routing. The gateway '
        'processes an average of 7,500 requests per second with a p95 latency of '
        '23 milliseconds. OAuth 2.0 with PKCE is enforced for all client applications, '
        'while internal service-to-service calls use mutual TLS with short-lived '
        'certificates rotated every 24 hours.'
    )

    # Figure 4
    img4_path = f'{WORKDIR}/fig_api_gateway.png'
    create_placeholder_image(img4_path, 'API Gateway Flow', 640, 420)
    doc.add_picture(img4_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_caption_paragraph(doc, 'Figure 4: API Gateway Flow')

    doc.add_paragraph(
        'Figure 4 depicts the request flow from client applications through the API '
        'gateway, including the authentication middleware, rate limiter, circuit breaker, '
        'and load balancer stages. Failed requests are routed to a dead-letter queue for '
        'later analysis and replay.'
    )

    # --- Section 6: Load Balancing ---
    doc.add_heading('6. Load Balancing Strategy', level=1)
    doc.add_paragraph(
        'Traffic distribution uses a two-tier load balancing approach. The global tier '
        'employs GeoDNS and anycast routing to direct users to the nearest regional '
        'endpoint. The regional tier uses weighted round-robin with health-check-aware '
        'routing to distribute requests across available pods.'
    )
    doc.add_paragraph(
        'During the Q3 2024 traffic spike event, the load balancing system successfully '
        'handled a 340% increase in traffic over baseline by dynamically scaling backend '
        'pools and redistributing connections to underutilized nodes. Zero downtime was '
        'recorded during the 6-hour event window.'
    )

    # Figure 5
    img5_path = f'{WORKDIR}/fig_load_balancer.png'
    create_placeholder_image(img5_path, 'Load Balancer Config', 640, 400)
    doc.add_picture(img5_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_caption_paragraph(doc, 'Figure 5: Load Balancer Configuration')

    doc.add_paragraph(
        'The load balancer configuration shown in Figure 5 includes health check intervals, '
        'connection draining parameters, and the weighted routing algorithm. The system '
        'supports both Layer 4 (TCP) and Layer 7 (HTTP) balancing modes.'
    )

    # --- Section 7: Monitoring ---
    doc.add_heading('7. Monitoring & Observability', level=1)
    doc.add_paragraph(
        'Comprehensive observability is achieved through a three-pillar approach: metrics, '
        'logs, and traces. Prometheus collects time-series metrics from all services with '
        '15-second scrape intervals. Structured JSON logs are shipped to a centralized '
        'Loki cluster with 90-day retention. Distributed traces using OpenTelemetry span '
        'all service boundaries, enabling end-to-end latency analysis.'
    )
    doc.add_paragraph(
        'Alerting rules are defined in code and deployed through the same CI/CD pipeline '
        'as application services. On-call engineers receive notifications through a '
        'multi-channel escalation policy that includes Slack, PagerDuty, and SMS fallback. '
        'Mean time to detection (MTTD) has improved from 8.3 minutes to 1.7 minutes since '
        'the observability platform was deployed.'
    )

    # Figure 6
    img6_path = f'{WORKDIR}/fig_monitoring.png'
    create_placeholder_image(img6_path, 'Monitoring Dashboard', 640, 430)
    doc.add_picture(img6_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_caption_paragraph(doc, 'Figure 6: Monitoring Dashboard')

    doc.add_paragraph(
        'Figure 6 shows the primary monitoring dashboard displaying real-time metrics for '
        'all three regions. Key performance indicators include request throughput, error '
        'rates, pod availability, and database connection pool utilization.'
    )

    doc.add_page_break()

    # --- Appendix ---
    doc.add_heading('Appendix A: Glossary', level=1)
    glossary_items = [
        ('VPC', 'Virtual Private Cloud — an isolated network segment within the cloud provider.'),
        ('gVisor', 'A lightweight container sandbox providing an additional security layer.'),
        ('PKCE', 'Proof Key for Code Exchange — an extension to OAuth 2.0 for public clients.'),
        ('GeoDNS', 'Geography-aware DNS resolution directing users to the nearest data center.'),
        ('MTTD', 'Mean Time To Detection — average time to identify an incident.'),
        ('p99', 'The 99th percentile response time — only 1% of requests are slower.'),
    ]
    for term, definition in glossary_items:
        p = doc.add_paragraph()
        run_term = p.add_run(f'{term}: ')
        run_term.bold = True
        p.add_run(definition)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
