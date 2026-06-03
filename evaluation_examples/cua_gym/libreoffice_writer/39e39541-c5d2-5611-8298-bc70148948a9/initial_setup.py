"""
Initial Setup: Legal contract with language set to None (no spell checking)
Task ID: writer_legal_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_language_none(element):
    """Set language to 'None' (zxx = no linguistic content) on a run or paragraph."""
    rPr = element.find(qn('w:rPr'))
    if rPr is None:
        rPr = element.makeelement(qn('w:rPr'), {})
        element.insert(0, rPr)
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = rPr.makeelement(qn('w:lang'), {})
        rPr.append(lang)
    lang.set(qn('w:val'), 'zxx')
    lang.set(qn('w:eastAsia'), 'zxx')
    lang.set(qn('w:bidi'), 'zxx')


def set_default_style_language_none(doc):
    """Set the Default Paragraph Style font language to 'None' (zxx)."""
    styles = doc.styles
    for style in styles:
        if style.name == 'Normal' or style.name == 'Default Paragraph Font':
            elem = style.element
            rPr = elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = elem.makeelement(qn('w:rPr'), {})
                elem.append(rPr)
            lang = rPr.find(qn('w:lang'))
            if lang is None:
                lang = rPr.makeelement(qn('w:lang'), {})
                rPr.append(lang)
            lang.set(qn('w:val'), 'zxx')
            lang.set(qn('w:eastAsia'), 'zxx')
            lang.set(qn('w:bidi'), 'zxx')


def create_initial():
    doc = Document()

    # Set default style language to None (zxx = no linguistic content)
    set_default_style_language_none(doc)

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        set_language_none(run._element)

    # --- Effective Date ---
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)
    run.font.italic = True
    set_language_none(run._element)

    doc.add_paragraph()  # blank line

    # --- Parties Section ---
    h = doc.add_heading('1. PARTIES', level=1)
    for run in h.runs:
        set_language_none(run._element)

    # Intentional misspelling: "agrrement" instead of "agreement"
    p = doc.add_paragraph()
    run = p.add_run(
        'This Professional Services Agrrement ("Agreement") is entered into as of the '
        'Effective Date by and between the following parties:'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run(
        'Pinnacle Dynamics Corporation, a Delaware corporation with its principal offices '
        'located at 4200 Innovation Boulevard, Suite 800, San Francisco, California 94105 '
        '(hereinafter referred to as "Service Provider").'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    # Intentional misspelling: "Massachuesetts" instead of "Massachusetts"
    run = p.add_run(
        'Horizon Ventures LLC, a limited liability company organized under the laws of '
        'Massachuesetts with its principal offices located at 1750 Beacon Street, '
        'Suite 300, Boston, Massachusetts 02134 (hereinafter referred to as "Client").'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    # --- Scope of Services ---
    h = doc.add_heading('2. SCOPE OF SERVICES', level=1)
    for run in h.runs:
        set_language_none(run._element)

    # Intentional misspelling: "responsibilites" instead of "responsibilities"
    p = doc.add_paragraph()
    run = p.add_run(
        'The Service Provider agrees to perform the following responsibilites and '
        'deliverables for the Client in accordance with the terms and conditions set '
        'forth herein:'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    items = [
        'Strategic technology consulting and advisory services for enterprise software modernization.',
        'Development and implementaton of cloud infrastructure migration plans.',  # "implementaton"
        'Quarterly performance assessments and comprehensive reporting dashboards.',
        'Staff training programs and knowledge transfer sesions for internal teams.',  # "sesions"
        'Ongoing technical support and maintanence during the contract period.',  # "maintanence"
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        set_language_none(run._element)

    # --- Compensation ---
    h = doc.add_heading('3. COMPENSATION', level=1)
    for run in h.runs:
        set_language_none(run._element)

    # Intentional misspelling: "recieve" instead of "receive"
    p = doc.add_paragraph()
    run = p.add_run(
        'In consideration of the services rendered, the Service Provider shall recieve '
        'compensation as follows:'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run(
        'a) A monthly retainer fee of Forty-Five Thousand Dollars ($45,000.00) payable '
        'on the first business day of each calendar month.'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    # Intentional misspelling: "exceding" instead of "exceeding"
    run = p.add_run(
        'b) Additional project-based fees for work exceding the scope defined in '
        'Section 2, billed at a rate of Three Hundred Fifty Dollars ($350.00) per hour.'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run(
        'c) Reimbursment of pre-approved travel and incidental expenses up to a maximum '
        'of Five Thousand Dollars ($5,000.00) per quarter.'
    )  # "Reimbursment"
    run.font.size = Pt(11)
    set_language_none(run._element)

    # --- Term and Termination ---
    h = doc.add_heading('4. TERM AND TERMINATION', level=1)
    for run in h.runs:
        set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run(
        'This Agreement shall commence on the Effective Date and shall remain in full '
        'force and effect for a period of twenty-four (24) months unless terminated '
        'earlier in acordance with this Section.'
    )  # "acordance"
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run(
        'Either party may terminate this Agreement with or without cause by providing '
        'ninety (90) days written notice to the other party. Upon termination, the '
        'Client shall pay the Service Provider for all services performd through the '
        'effective date of termination.'
    )  # "performd"
    run.font.size = Pt(11)
    set_language_none(run._element)

    # --- Confidentiality ---
    h = doc.add_heading('5. CONFIDENTIALITY', level=1)
    for run in h.runs:
        set_language_none(run._element)

    p = doc.add_paragraph()
    # Intentional misspelling: "confidental" instead of "confidential"
    run = p.add_run(
        'Each party acknowledges that during the course of this Agreement, it may have '
        'access to confidental and proprietary information belonging to the other party. '
        'Both parties agree to maintain the confidentiality of such information and to '
        'not disclose it to any third party without prior written consent.'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    # --- Governing Law ---
    h = doc.add_heading('6. GOVERNING LAW', level=1)
    for run in h.runs:
        set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run(
        'This Agreement shall be governed by and construed in accordance with the laws '
        'of the State of California, without regard to its conflict of laws provisions. '
        'Any disputes arising under this Agreement shall be resolved through binding '
        'arbitration in San Francisco, California.'
    )
    run.font.size = Pt(11)
    set_language_none(run._element)

    # --- Signatures ---
    doc.add_paragraph()
    h = doc.add_heading('SIGNATURES', level=1)
    for run in h.runs:
        set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run('FOR SERVICE PROVIDER:')
    run.bold = True
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run('Name: ____________________________')
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run('Title: ____________________________')
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run('Date: ____________________________')
    run.font.size = Pt(11)
    set_language_none(run._element)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('FOR CLIENT:')
    run.bold = True
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run('Name: ____________________________')
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run('Title: ____________________________')
    run.font.size = Pt(11)
    set_language_none(run._element)

    p = doc.add_paragraph()
    run = p.add_run('Date: ____________________________')
    run.font.size = Pt(11)
    set_language_none(run._element)

    # Also set language on document-level default run properties
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    # Set document-level default language via docDefaults in styles
    styles_elem = doc.styles.element
    docDefaults = styles_elem.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = styles_elem.makeelement(qn('w:docDefaults'), {})
        styles_elem.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = docDefaults.makeelement(qn('w:rPrDefault'), {})
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = rPrDefault.makeelement(qn('w:rPr'), {})
        rPrDefault.append(rPr)
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = rPr.makeelement(qn('w:lang'), {})
        rPr.append(lang)
    lang.set(qn('w:val'), 'zxx')
    lang.set(qn('w:eastAsia'), 'zxx')
    lang.set(qn('w:bidi'), 'zxx')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
