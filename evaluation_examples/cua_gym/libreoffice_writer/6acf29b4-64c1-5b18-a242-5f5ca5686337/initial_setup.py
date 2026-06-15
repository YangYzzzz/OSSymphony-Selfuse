"""
Initial Setup: Project management notes document without Emphasis style applied.
Task ID: writer_txtfmt_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_016'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/pm_notes.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set default font to Liberation Serif 12pt for the document
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style.font.size = Pt(12)

    # Heading: Project Management Fundamentals
    heading = doc.add_heading('Project Management Fundamentals', level=1)
    for run in heading.runs:
        run.font.name = 'Liberation Serif'

    # Paragraph 1: Introduction
    para1 = doc.add_paragraph()
    run1 = para1.add_run(
        'Project management is the discipline of initiating, planning, executing, '
        'controlling, and closing the work of a team to achieve specific goals and '
        'meet specific success criteria at a specified time. Effective project '
        'management requires clear communication, structured processes, and skilled '
        'coordination among all stakeholders involved in the effort.'
    )
    run1.font.name = 'Liberation Serif'
    run1.font.size = Pt(12)

    # Paragraph 2: Contains 'critical path analysis' (NOT emphasized in initial state)
    para2 = doc.add_paragraph()
    run2 = para2.add_run(
        'The most important technique in scheduling is critical path analysis, '
        'which determines the longest sequence of dependent tasks. By identifying '
        'the critical path, project managers can focus their attention on the '
        'activities that directly affect the project completion date. Any delay '
        'along the critical path results in an equivalent delay to the overall '
        'project timeline.'
    )
    run2.font.name = 'Liberation Serif'
    run2.font.size = Pt(12)

    # Paragraph 3: Conclusion
    para3 = doc.add_paragraph()
    run3 = para3.add_run(
        'Resource allocation and risk management are equally vital components of '
        'a successful project. Project managers must continuously monitor progress '
        'against the baseline schedule, identify potential bottlenecks, and '
        'implement corrective actions when variances exceed acceptable thresholds. '
        'Regular status reviews and stakeholder updates ensure alignment and '
        'maintain project momentum throughout the execution phase.'
    )
    run3.font.name = 'Liberation Serif'
    run3.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
