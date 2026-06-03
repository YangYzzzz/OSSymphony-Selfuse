"""
Reward Script: Apply 'Emphasis' character style to 'critical path analysis' in paragraph 2
Task ID: writer_txtfmt_016
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): The phrase 'critical path analysis' in paragraph 2 has the built-in
                      'Emphasis' character style applied (w:rStyle val='Emphasis').
                      FAILS on initial (single run, no rStyle) → PASSES on golden.
  Component 2 (0.4): Paragraph 2 is split into multiple runs (>1) AND all text is intact.
                      Initial state: 1 run → FAILS. Golden state: 3 runs → PASSES.
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_016'
FILE_PATH = f'{WORKDIR}/Desktop/pm_notes.docx'

TARGET_PHRASE = 'critical path analysis'
TARGET_PARA_INDEX = 2  # 0-indexed: heading=0, para1=1, para2=2, para3=3

# Expected full text of paragraph 2 (from task context)
EXPECTED_PARA2_TEXT = (
    'The most important technique in scheduling is critical path analysis, '
    'which determines the longest sequence of dependent tasks. By identifying '
    'the critical path, project managers can focus their attention on the '
    'activities that directly affect the project completion date. Any delay '
    'along the critical path results in an equivalent delay to the overall '
    'project timeline.'
)


def get_run_style_val(run):
    """Return the w:rStyle val attribute from a run's rPr, or None if not present."""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        return None
    rStyle = rPr.find(qn('w:rStyle'))
    if rStyle is None:
        return None
    return rStyle.get(qn('w:val'))


def find_emphasis_run_for_phrase(para, phrase):
    """
    Search para.runs for a run whose text equals phrase exactly AND
    has the Emphasis character style applied (via rStyle XML attribute
    or run.style.name).
    Returns the matching run, or None.
    """
    for run in para.runs:
        style_val = get_run_style_val(run)
        has_emphasis = (style_val == 'Emphasis') or (
            run.style is not None and run.style.name == 'Emphasis'
        )
        if has_emphasis and run.text == phrase:
            return run
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: load the file
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have at least 4 paragraphs
    if len(doc.paragraphs) < 4:
        print(f"CRITICAL: Expected at least 4 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: paragraph 2 text must still contain the target phrase
    para2 = doc.paragraphs[TARGET_PARA_INDEX]
    if TARGET_PHRASE not in para2.text:
        print(f"CRITICAL: Paragraph 2 does not contain '{TARGET_PHRASE}'. "
              f"Text may have been corrupted.")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: The phrase 'critical path analysis' must exist as a distinct run
    #              in paragraph 2 with the 'Emphasis' character style (0.6 points)
    #              This FAILS on initial (single run, no rStyle) and
    #              PASSES on golden (run with rStyle val='Emphasis').
    try:
        matched_run = find_emphasis_run_for_phrase(para2, TARGET_PHRASE)

        if matched_run is not None:
            print(f"PASS: Component 1 — '{TARGET_PHRASE}' has 'Emphasis' character style "
                  f"in paragraph 2 (rStyle=Emphasis confirmed) (0.6 pts)")
            total_score += 0.6
        else:
            # Distinguish: was there an Emphasis run at all, or was there no Emphasis run?
            emphasis_runs = [r for r in para2.runs
                             if get_run_style_val(r) == 'Emphasis' or
                             (r.style is not None and r.style.name == 'Emphasis')]
            if not emphasis_runs:
                print(f"FAIL: Component 1 — No run with 'Emphasis' character style found "
                      f"in paragraph 2 (initial state has no rStyle set)")
            else:
                texts = [r.text for r in emphasis_runs]
                print(f"FAIL: Component 1 — Emphasis style found but not on exact phrase "
                      f"'{TARGET_PHRASE}'. Emphasis runs: {texts}")

    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Paragraph 2 must have been split into multiple runs (>1) to isolate
    #              the target phrase, AND the full text must remain intact (0.4 points).
    #              Initial state: single run (1 run) → FAILS this check.
    #              Golden state: 3 runs ('...is ', 'critical path analysis', ', which...') → PASSES.
    try:
        num_runs = len(para2.runs)
        actual_text = para2.text.strip()
        expected_text = EXPECTED_PARA2_TEXT.strip()
        text_intact = (actual_text == expected_text)

        if num_runs > 1 and text_intact:
            print(f"PASS: Component 2 — Paragraph 2 split into {num_runs} runs and "
                  f"text content is fully preserved (0.4 pts)")
            total_score += 0.4
        elif num_runs <= 1:
            print(f"FAIL: Component 2 — Paragraph 2 has only {num_runs} run(s); "
                  f"expected >1 runs indicating phrase was isolated for styling")
        else:
            print(f"FAIL: Component 2 — Paragraph 2 has {num_runs} runs but text content "
                  f"differs from expected. Lengths: actual={len(actual_text)}, "
                  f"expected={len(expected_text)}")

    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: test against the canonical artifact path
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
