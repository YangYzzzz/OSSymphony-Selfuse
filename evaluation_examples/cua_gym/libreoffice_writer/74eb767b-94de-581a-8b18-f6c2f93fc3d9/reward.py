"""
Reward Script: Format results table to APA style
Task ID: writer_acad_028
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35) — Vertical borders removed (left/right on all cells are none)
  Component 2 (0.35) — Three horizontal lines: above header, below header, below last row
  Component 3 (0.30) — Columns 2-4 center-aligned, column 1 left-aligned
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_028'


def get_cell_border(cell, side):
    """
    Return the border 'val' attribute for a given side of a cell.
    Returns None if no cell-level border is defined for that side.
    A value of 'none' means explicitly no border; 'single' means a visible line.
    """
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return None
    elem = tcBorders.find(qn('w:' + side))
    if elem is None:
        return None
    return elem.get(qn('w:val'))


def get_table_border(table, side):
    """
    Return the table-level border 'val' for a given side (top, bottom, left, right, insideH, insideV).
    """
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        return None
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        return None
    elem = tblBorders.find(qn('w:' + side))
    if elem is None:
        return None
    return elem.get(qn('w:val'))


def border_is_visible(val):
    """Check if a border value represents a visible border."""
    if val is None:
        return False  # not explicitly set at this level
    return val not in ('none', 'nil', '')


def verify_task(file_path):
    """
    Verify APA table formatting with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least one table with 6 rows and 4 columns
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    if num_rows < 2 or num_cols != 4:
        print(f"FAIL: Table dimensions unexpected: {num_rows} rows x {num_cols} cols")
        print("REWARD: 0.0")
        return 0.0

    last_row = num_rows - 1

    # ======================================================================
    # Component 1: Vertical borders removed (0.35 points)
    # APA style requires NO vertical borders at all.
    # Check: table-level insideV is 'none', and all cells have left/right = 'none' or not set.
    # This should FAIL on initial (Table Grid has vertical borders by style) and PASS on golden.
    # ======================================================================
    try:
        # Strategy: The initial file uses "Table Grid" style which provides vertical borders
        # via the style definition. The golden file either:
        # (a) sets table-level insideV/left/right to 'none', or
        # (b) sets cell-level left/right to 'none' on all cells
        # We check that no cell has a visible vertical border.

        # First check table-level: insideV, left, right should be 'none' if present
        tbl_insideV = get_table_border(table, 'insideV')
        tbl_left = get_table_border(table, 'left')
        tbl_right = get_table_border(table, 'right')

        # If table-level borders are explicitly set to 'none', that's a strong signal
        table_level_vert_removed = (
            tbl_insideV is not None and not border_is_visible(tbl_insideV) and
            tbl_left is not None and not border_is_visible(tbl_left) and
            tbl_right is not None and not border_is_visible(tbl_right)
        )

        # Also check cell-level: all cells should have left/right = 'none' or rely on table-level 'none'
        cell_vert_ok = True
        for ri in range(num_rows):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                cell_left = get_cell_border(cell, 'left')
                cell_right = get_cell_border(cell, 'right')
                if border_is_visible(cell_left) or border_is_visible(cell_right):
                    cell_vert_ok = False
                    break
            if not cell_vert_ok:
                break

        # The key change: the initial file has NO explicit border overrides (relies on Table Grid style
        # which has visible borders). The golden file explicitly sets them to 'none'.
        # So we require either table-level vertical 'none' OR all cells explicitly set to 'none'.
        vert_removed = table_level_vert_removed and cell_vert_ok

        if vert_removed:
            print(f"PASS: Component 1 — All vertical borders removed (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — Vertical borders still present. "
                  f"table_level_vert_removed={table_level_vert_removed}, cell_vert_ok={cell_vert_ok}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ======================================================================
    # Component 2: Three horizontal lines only (0.35 points)
    # APA style: lines above header row, below header row, below last data row.
    # Check cell-level borders for the pattern:
    #   - Row 0 cells: top=single, bottom=single
    #   - Rows 1 to last_row-1: top=none, bottom=none (no horizontal lines in middle)
    #   - Row last_row: bottom=single (and top=none)
    # This FAILS on initial (Table Grid has all borders via style) and PASSES on golden.
    # ======================================================================
    try:
        horiz_checks_passed = 0
        horiz_checks_total = 3

        # Sub-check 2a: Header row top border (line above header)
        header_top_ok = True
        for ci in range(num_cols):
            val = get_cell_border(table.cell(0, ci), 'top')
            if not border_is_visible(val):
                header_top_ok = False
                break

        if header_top_ok:
            print("PASS: Component 2a — Line above header row present")
            horiz_checks_passed += 1
        else:
            print("FAIL: Component 2a — Line above header row missing")

        # Sub-check 2b: Header row bottom border (line below header)
        header_bottom_ok = True
        for ci in range(num_cols):
            val = get_cell_border(table.cell(0, ci), 'bottom')
            if not border_is_visible(val):
                header_bottom_ok = False
                break

        if header_bottom_ok:
            print("PASS: Component 2b — Line below header row present")
            horiz_checks_passed += 1
        else:
            print("FAIL: Component 2b — Line below header row missing")

        # Sub-check 2c: Last row bottom border (line below last row)
        # AND middle rows have no visible horizontal borders (no extra lines)
        last_bottom_ok = True
        for ci in range(num_cols):
            val = get_cell_border(table.cell(last_row, ci), 'bottom')
            if not border_is_visible(val):
                last_bottom_ok = False
                break

        # Also verify that middle data rows (1 to last_row-1) do NOT have bottom borders
        # and rows 1 to last_row do NOT have top borders (to ensure only 3 lines)
        middle_clean = True
        for ri in range(1, num_rows):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                top_val = get_cell_border(cell, 'top')
                if border_is_visible(top_val):
                    middle_clean = False
                    break
            if not middle_clean:
                break

        for ri in range(1, last_row):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                bottom_val = get_cell_border(cell, 'bottom')
                if border_is_visible(bottom_val):
                    middle_clean = False
                    break
            if not middle_clean:
                break

        if last_bottom_ok and middle_clean:
            print("PASS: Component 2c — Line below last row present, no extra horizontal lines")
            horiz_checks_passed += 1
        else:
            print(f"FAIL: Component 2c — last_bottom_ok={last_bottom_ok}, middle_clean={middle_clean}")

        # Award points proportionally
        comp2_score = 0.35 * (horiz_checks_passed / horiz_checks_total)
        if comp2_score > 0:
            print(f"PASS: Component 2 — Horizontal lines: {horiz_checks_passed}/{horiz_checks_total} checks ({comp2_score:.3f} pts)")
            total_score += comp2_score
        else:
            print(f"FAIL: Component 2 — No horizontal line checks passed")

    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ======================================================================
    # Component 3: Column alignment (0.30 points)
    # Columns 2-4 (index 1-3) must be CENTER-aligned.
    # Column 1 (index 0) must be LEFT-aligned (None or LEFT).
    # This FAILS on initial (all None/LEFT) and PASSES on golden.
    # ======================================================================
    try:
        center_correct = 0
        center_total = 0

        for ri in range(num_rows):
            for ci in range(1, 4):  # columns 1, 2, 3 (the numeric columns)
                center_total += 1
                cell = table.cell(ri, ci)
                for p in cell.paragraphs:
                    align = p.paragraph_format.alignment
                    if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        center_correct += 1
                    break  # only check first paragraph in cell

        # Column 0 should NOT be center-aligned (should be left/None)
        col0_ok = True
        for ri in range(num_rows):
            cell = table.cell(ri, 0)
            for p in cell.paragraphs:
                align = p.paragraph_format.alignment
                if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    col0_ok = False
                break

        if center_correct == center_total and col0_ok:
            print(f"PASS: Component 3 — All numeric columns center-aligned, col 1 left-aligned (0.30 pts)")
            total_score += 0.30
        elif center_correct > 0 and col0_ok:
            partial = 0.30 * (center_correct / center_total)
            print(f"PARTIAL: Component 3 — {center_correct}/{center_total} numeric cells centered ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — center_correct={center_correct}/{center_total}, col0_ok={col0_ok}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
