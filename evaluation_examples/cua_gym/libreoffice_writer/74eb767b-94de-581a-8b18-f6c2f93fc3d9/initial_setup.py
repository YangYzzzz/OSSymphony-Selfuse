"""
Initial Setup: APA-style table formatting task
Task ID: writer_acad_028
Domain: libreoffice_writer
Creates a Writer document with a 6-row, 4-column table with full grid borders.
All columns default left-aligned. No APA formatting applied.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_028'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell. Each border arg is a dict with sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, props in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if props is not None:
            el = OxmlElement(f'w:{edge}')
            for k, v in props.items():
                el.set(qn(f'w:{k}'), str(v))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Cognitive Performance Under Sleep Deprivation', level=1)

    # Intro paragraph
    doc.add_paragraph(
        'Table 1 presents the descriptive statistics for cognitive performance '
        'measures across both experimental groups. Participants were assessed '
        'using standardized instruments after a 36-hour protocol.'
    )

    # Create 6-row, 4-column table with full grid borders
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'  # Full grid borders

    # Headers
    headers = ['Measure', 'Group A', 'Group B', 'Effect Size']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # Data rows - realistic cognitive performance data
    data = [
        ['Reaction Time (ms)',    '342.7 (28.4)',  '298.1 (22.6)',  '0.83'],
        ['Working Memory',        '5.2 (1.8)',     '7.4 (1.3)',     '1.41'],
        ['Sustained Attention',   '78.3 (9.1)',    '91.6 (5.7)',    '1.75'],
        ['Executive Function',    '12.4 (3.6)',    '16.8 (2.9)',    '1.35'],
        ['Processing Speed',      '45.1 (6.2)',    '52.9 (4.8)',    '1.41'],
    ]

    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Note paragraph below table
    doc.add_paragraph('')
    note = doc.add_paragraph()
    run = note.add_run('Note.')
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run2 = note.add_run(' Standard deviations are shown in parentheses. Effect sizes are Cohen\'s d.')
    run2.font.size = Pt(10)
    run2.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
