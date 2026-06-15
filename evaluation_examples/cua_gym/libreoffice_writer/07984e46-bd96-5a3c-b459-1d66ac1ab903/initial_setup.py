"""
Initial Setup: Create a Writer document with eight footnotes using default Footnote Anchor style.
Task ID: writer_bs_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_039'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def xml_escape(text):
    """Escape special XML characters."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&apos;')


def add_footnote(paragraph, footnote_text, footnote_id):
    """Add a footnote to a paragraph using raw XML manipulation."""
    doc_part = paragraph.part

    footnotes_part = None
    for rel in doc_part.rels.values():
        if 'footnotes' in rel.reltype:
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        ct = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
        footnotes_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )
        footnotes_part = Part(
            partname=PackURI('/word/footnotes.xml'),
            content_type=ct,
            blob=footnotes_xml.encode('utf-8'),
            package=doc_part.package,
        )
        doc_part.relate_to(footnotes_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')

    footnotes_elem = etree.fromstring(footnotes_part.blob)

    safe_text = xml_escape(footnote_text)
    fn_xml = (
        f'<w:footnote xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:id="{footnote_id}">'
        f'<w:p>'
        f'<w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
        f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'<w:footnoteRef/></w:r>'
        f'<w:r><w:t xml:space="preserve"> {safe_text}</w:t></w:r>'
        f'</w:p>'
        f'</w:footnote>'
    )
    fn_elem = etree.fromstring(fn_xml)
    footnotes_elem.append(fn_elem)

    footnotes_part._blob = etree.tostring(footnotes_elem, xml_declaration=True, encoding='UTF-8', standalone=True)

    ref_xml = (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'<w:footnoteReference w:id="{footnote_id}"/>'
        f'</w:r>'
    )
    ref_elem = etree.fromstring(ref_xml)
    paragraph._element.append(ref_elem)


def ensure_footnote_styles(doc):
    """Ensure 'Footnote Text' paragraph style and 'Footnote Reference' character style exist."""
    styles_elem = doc.styles.element

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Check if FootnoteReference style exists
    existing = styles_elem.findall('.//w:style', nsmap)
    has_fn_ref = False
    has_fn_text = False
    for s in existing:
        sid = s.get(qn('w:styleId'))
        if sid == 'FootnoteReference':
            has_fn_ref = True
        if sid == 'FootnoteText':
            has_fn_text = True

    if not has_fn_ref:
        # Create FootnoteReference character style (default: superscript only)
        style_xml = (
            '<w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:type="character" w:styleId="FootnoteReference">'
            '<w:name w:val="Footnote Reference"/>'
            '<w:basedOn w:val="DefaultParagraphFont"/>'
            '<w:rPr>'
            '<w:vertAlign w:val="superscript"/>'
            '</w:rPr>'
            '</w:style>'
        )
        styles_elem.append(etree.fromstring(style_xml))

    if not has_fn_text:
        # Create FootnoteText paragraph style
        style_xml = (
            '<w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:type="paragraph" w:styleId="FootnoteText">'
            '<w:name w:val="Footnote Text"/>'
            '<w:basedOn w:val="Normal"/>'
            '<w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:rPr><w:sz w:val="20"/></w:rPr>'
            '</w:style>'
        )
        styles_elem.append(etree.fromstring(style_xml))


def create_initial():
    doc = Document()

    # Ensure footnote-related styles exist
    ensure_footnote_styles(doc)

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('The Impact of Renewable Energy on Global Economic Development', level=1)

    # --- Author line ---
    author = doc.add_paragraph()
    author.paragraph_format.space_after = Pt(12)
    run = author.add_run('Dr. Elena Rodriguez, Institute for Sustainable Policy Research')
    run.italic = True
    run.font.size = Pt(10)

    # --- Abstract ---
    doc.add_heading('Abstract', level=2)
    abstract = doc.add_paragraph(
        'The transition to renewable energy sources has been one of the most significant economic '
        'shifts of the 21st century. This paper examines the multi-faceted impact of renewable energy '
        'adoption on global economic development, employment patterns, and trade relationships.'
    )

    # --- Introduction with footnotes 1 and 2 ---
    doc.add_heading('1. Introduction', level=2)

    p1 = doc.add_paragraph(
        'The global energy landscape has undergone a dramatic transformation over the past two decades. '
        'According to the International Renewable Energy Agency, total renewable power capacity reached '
        '3,372 gigawatts by the end of 2023'
    )
    add_footnote(p1, 'IRENA (2024). Renewable Capacity Statistics 2024. Abu Dhabi: International Renewable Energy Agency.', 1)
    run = p1.add_run(
        '. This represents a compound annual growth rate of 8.3% since 2010, far exceeding '
        'the growth rate of fossil fuel-based generation'
    )
    add_footnote(p1, 'Global Energy Monitor (2023). Annual Report on Fossil Fuel Phase-Out Progress, pp. 42-58.', 2)
    p1.add_run('.')

    # --- Economic Impact section with footnotes 3 and 4 ---
    doc.add_heading('2. Economic Impact Assessment', level=2)

    p2 = doc.add_paragraph(
        'Investment in renewable energy infrastructure has created substantial economic multiplier effects. '
        'A comprehensive study by the World Bank estimated that every dollar invested in clean energy '
        'generates approximately $3.20 in local economic activity'
    )
    add_footnote(p2, 'World Bank Group (2023). The Economic Multiplier Effects of Clean Energy Investment. Working Paper No. 9847.', 3)
    p2.add_run(
        '. Furthermore, the renewable energy sector employed an estimated 13.7 million people worldwide '
        'in 2023, a 15% increase from the previous year'
    )
    add_footnote(p2, 'IRENA & ILO (2024). Renewable Energy and Jobs: Annual Review 2024, Geneva.', 4)
    p2.add_run('.')

    # --- Employment section with footnotes 5 and 6 ---
    doc.add_heading('3. Employment and Workforce Transition', level=2)

    p3 = doc.add_paragraph(
        'The shift toward renewable energy has significant implications for labor markets globally. '
        'Solar photovoltaic installation alone accounted for 4.9 million jobs in 2023, making it the '
        'largest employer in the renewable energy sector'
    )
    add_footnote(p3, 'Solar Energy Industries Association (2024). National Solar Jobs Census 2024, pp. 12-29.', 5)
    p3.add_run(
        '. Regional analysis reveals that developing nations in Southeast Asia and Sub-Saharan Africa '
        'experienced the most rapid growth in green employment, with year-over-year increases averaging 22%'
    )
    add_footnote(p3, 'United Nations Environment Programme (2023). Green Economy Progress Report: Employment Trends, Nairobi.', 6)
    p3.add_run('.')

    # --- Trade Relationships section with footnotes 7 and 8 ---
    doc.add_heading('4. International Trade Relationships', level=2)

    p4 = doc.add_paragraph(
        'The transition to renewable energy has reshaped international trade dynamics. Countries that '
        'were previously dependent on fossil fuel imports have significantly reduced their trade deficits. '
        'For instance, the European Union reduced its energy import bill by approximately EUR 47 billion '
        'between 2019 and 2023'
    )
    add_footnote(p4, 'European Commission (2024). EU Energy in Figures: Statistical Pocketbook 2024, Brussels.', 7)
    p4.add_run(
        '. Simultaneously, new trade patterns have emerged around critical minerals and manufacturing '
        'supply chains essential for renewable energy technologies'
    )
    add_footnote(p4, 'International Energy Agency (2023). Critical Minerals Market Review 2023, Paris: OECD/IEA.', 8)
    p4.add_run('.')

    # --- Conclusion ---
    doc.add_heading('5. Conclusion', level=2)
    doc.add_paragraph(
        'The evidence presented in this paper demonstrates that the transition to renewable energy is not '
        'merely an environmental imperative but also a significant economic opportunity. The creation of '
        'millions of new jobs, the reduction in trade deficits for energy-importing nations, and the '
        'substantial multiplier effects of clean energy investment all point to a future where economic '
        'growth and environmental sustainability are mutually reinforcing.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
