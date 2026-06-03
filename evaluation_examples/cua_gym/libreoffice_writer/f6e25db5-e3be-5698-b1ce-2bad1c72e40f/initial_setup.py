"""
Initial Setup: Apply paragraph styles to a legal Court Brief document.
Task ID: writer_pd_026
Domain: libreoffice_writer

Creates a 15-page legal court brief where all text uses Default Paragraph Style
in 12pt Calibri with single spacing. Block quotes are indented but not styled
differently. Headings are manually bold but use the default paragraph style.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_paragraph(doc, text):
    """Add a body paragraph in 12pt Calibri, single spacing, no indent."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(6)
    return para


def add_heading_paragraph(doc, text, level=1):
    """Add a heading as default paragraph style with manual bold - NOT using Heading styles."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_block_quote(doc, text):
    """Add a block quote indented but using same default style (not italic, not 11pt)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    para.paragraph_format.left_indent = Cm(2)
    para.paragraph_format.right_indent = Cm(2)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(6)
    return para


def create_initial():
    doc = Document()

    # Page setup: US Letter with standard margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # =============================================
    # COVER PAGE (Page 1)
    # =============================================
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "")
    p = add_body_paragraph(doc, "IN THE UNITED STATES DISTRICT COURT")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = add_body_paragraph(doc, "FOR THE NORTHERN DISTRICT OF CALIFORNIA")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_body_paragraph(doc, "")

    p = add_body_paragraph(doc, "MERIDIAN TECHNOLOGIES, INC.,")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = add_body_paragraph(doc, "Plaintiff,")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = add_body_paragraph(doc, "v.")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = add_body_paragraph(doc, "CORNERSTONE DATA SOLUTIONS, LLC,")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = add_body_paragraph(doc, "Defendant.")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_body_paragraph(doc, "")

    p = add_body_paragraph(doc, "Case No. 3:2025-cv-04817-JST")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_body_paragraph(doc, "")

    p = doc.add_paragraph()
    run = p.add_run("PLAINTIFF'S MEMORANDUM OF POINTS AND AUTHORITIES")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph()
    run = p.add_run("IN SUPPORT OF MOTION FOR PRELIMINARY INJUNCTION")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing = 1.0

    # Page break after cover
    doc.add_page_break()

    # =============================================
    # TABLE OF CONTENTS (Page 2)
    # =============================================
    add_heading_paragraph(doc, "TABLE OF CONTENTS")

    toc_items = [
        ("I.", "INTRODUCTION", "1"),
        ("II.", "STATEMENT OF FACTS", "3"),
        ("", "A. Background of the Parties", "3"),
        ("", "B. The Contractual Relationship", "5"),
        ("", "C. Defendant's Breach and Misappropriation", "6"),
        ("III.", "LEGAL STANDARD FOR PRELIMINARY INJUNCTION", "8"),
        ("IV.", "ARGUMENT", "9"),
        ("", "A. Likelihood of Success on the Merits", "9"),
        ("", "B. Irreparable Harm", "11"),
        ("", "C. Balance of Equities", "12"),
        ("", "D. Public Interest", "13"),
        ("V.", "CONCLUSION", "14"),
    ]

    for num, title, page in toc_items:
        text = f"{num} {title}{'.' * (60 - len(num) - len(title))} {page}" if num else f"    {title}{'.' * (56 - len(title))} {page}"
        add_body_paragraph(doc, text)

    doc.add_page_break()

    # =============================================
    # TABLE OF AUTHORITIES (Page 3)
    # =============================================
    add_heading_paragraph(doc, "TABLE OF AUTHORITIES")

    add_heading_paragraph(doc, "Cases")
    cases = [
        "Winter v. Natural Resources Defense Council, Inc., 555 U.S. 7 (2008)",
        "eBay Inc. v. MercExchange, LLC, 547 U.S. 388 (2006)",
        "Cadence Design Systems, Inc. v. Avant! Corp., 125 F.3d 824 (9th Cir. 1997)",
        "MAI Systems Corp. v. Peak Computer, Inc., 991 F.2d 511 (9th Cir. 1993)",
        "Apple Inc. v. Samsung Electronics Co., 735 F.3d 1352 (Fed. Cir. 2013)",
        "Ruckelshaus v. Monsanto Co., 467 U.S. 986 (1984)",
        "Silvaco Data Systems v. Intel Corp., 184 Cal. App. 4th 210 (2010)",
        "DVD Copy Control Ass'n v. Bunner, 116 Cal. App. 4th 241 (2004)",
    ]
    for c in cases:
        add_body_paragraph(doc, c)

    add_heading_paragraph(doc, "Statutes")
    statutes = [
        "18 U.S.C. § 1836 (Defend Trade Secrets Act)",
        "Cal. Civ. Code § 3426 et seq. (California Uniform Trade Secrets Act)",
        "17 U.S.C. § 101 et seq. (Copyright Act of 1976)",
    ]
    for s in statutes:
        add_body_paragraph(doc, s)

    doc.add_page_break()

    # =============================================
    # I. INTRODUCTION (Pages 4-5)
    # =============================================
    add_heading_paragraph(doc, "I. INTRODUCTION")

    add_body_paragraph(doc, (
        "Plaintiff Meridian Technologies, Inc. ('Meridian') respectfully submits this memorandum "
        "in support of its motion for a preliminary injunction against Defendant Cornerstone Data "
        "Solutions, LLC ('Cornerstone'). This case arises from Cornerstone's systematic and willful "
        "misappropriation of Meridian's proprietary trade secrets, including its core machine learning "
        "algorithms, customer relationship databases, and confidential pricing models."
    ))

    add_body_paragraph(doc, (
        "Meridian is a leader in enterprise data analytics, serving over 2,300 corporate clients "
        "across North America, Europe, and Asia-Pacific. Founded in 2008 by Dr. Elena Vasquez and "
        "Dr. James Thornton, Meridian has invested more than $147 million in research and development "
        "over the past decade, resulting in 23 registered patents and a proprietary software platform "
        "known internally as 'Nexus Analytics Suite.'"
    ))

    add_body_paragraph(doc, (
        "In March 2024, three senior engineers—Dr. Priya Ramaswamy, Kevin O'Brien, and Lisa "
        "Park—departed Meridian and joined Cornerstone within days of each other. Within six months "
        "of their departure, Cornerstone launched 'DataForge Pro,' a competing analytics platform "
        "that bears striking resemblance to Meridian's proprietary Nexus system in its architecture, "
        "algorithms, and user interface design."
    ))

    add_body_paragraph(doc, (
        "As demonstrated below, Meridian is entitled to a preliminary injunction because it satisfies "
        "all four factors of the Winter test: (1) likelihood of success on the merits; (2) irreparable "
        "harm absent injunctive relief; (3) the balance of equities tips in Meridian's favor; and "
        "(4) an injunction serves the public interest."
    ))

    # Block quote - legal citation
    add_block_quote(doc, (
        "\"A plaintiff seeking a preliminary injunction must establish that he is likely to succeed "
        "on the merits, that he is likely to suffer irreparable harm in the absence of preliminary "
        "relief, that the balance of equities tips in his favor, and that an injunction is in the "
        "public interest.\" Winter v. Natural Resources Defense Council, Inc., 555 U.S. 7, 20 (2008)."
    ))

    add_body_paragraph(doc, (
        "The evidence overwhelmingly supports each of these elements, and this Court should grant "
        "Meridian's motion to prevent further irreparable harm to its competitive position and "
        "proprietary interests."
    ))

    doc.add_page_break()

    # =============================================
    # II. STATEMENT OF FACTS (Pages 5-8)
    # =============================================
    add_heading_paragraph(doc, "II. STATEMENT OF FACTS")

    add_heading_paragraph(doc, "A. Background of the Parties")

    add_body_paragraph(doc, (
        "Meridian Technologies, Inc. is a Delaware corporation headquartered in San Francisco, "
        "California. Since its founding in 2008, Meridian has grown from a five-person startup to "
        "an organization employing 1,847 individuals across twelve offices worldwide. The company's "
        "annual revenue for fiscal year 2024 was approximately $623 million, representing a compound "
        "annual growth rate of 34% over the preceding five years."
    ))

    add_body_paragraph(doc, (
        "Meridian's flagship product, the Nexus Analytics Suite, is a comprehensive enterprise data "
        "analytics platform that integrates machine learning algorithms, natural language processing "
        "capabilities, and advanced visualization tools. The development of the Nexus platform has "
        "required sustained investment exceeding $147 million, with over 340 engineers and data "
        "scientists contributing to its creation over a period of nine years."
    ))

    add_body_paragraph(doc, (
        "Cornerstone Data Solutions, LLC is a Texas limited liability company formed in January 2023. "
        "Cornerstone was founded by former IBM executive Robert Harrington and venture capital firm "
        "Pinnacle Partners. Despite being a relatively new entrant to the enterprise analytics market, "
        "Cornerstone has rapidly gained market share following the launch of its DataForge Pro platform "
        "in September 2024."
    ))

    add_heading_paragraph(doc, "B. The Contractual Relationship")

    add_body_paragraph(doc, (
        "Each of the three departing engineers signed comprehensive employment agreements upon joining "
        "Meridian. These agreements included non-disclosure provisions, invention assignment clauses, "
        "non-solicitation covenants, and non-compete restrictions. Specifically, Section 7.2 of each "
        "agreement provides:"
    ))

    add_block_quote(doc, (
        "\"Employee agrees that during the term of employment and for a period of twenty-four (24) "
        "months following termination for any reason, Employee shall not directly or indirectly "
        "disclose, use, or exploit any Confidential Information or Trade Secrets of the Company for "
        "any purpose other than the performance of Employee's duties. This obligation survives "
        "termination of employment regardless of the circumstances of such termination.\""
    ))

    add_body_paragraph(doc, (
        "Furthermore, Section 8.1 of the employment agreements contains an invention assignment "
        "clause requiring all work product created during the course of employment to be assigned "
        "to Meridian. Dr. Ramaswamy, Mr. O'Brien, and Ms. Park each acknowledged receipt of and "
        "agreement to these terms in writing upon their hiring."
    ))

    add_body_paragraph(doc, (
        "Dr. Priya Ramaswamy joined Meridian in 2015 as a Senior Machine Learning Engineer and was "
        "promoted to Principal Architect in 2020. During her tenure, she served as the lead developer "
        "of Nexus's core predictive analytics engine, designated internally as 'Project Phoenix.' "
        "Kevin O'Brien served as Director of Client Solutions from 2017 to 2024, managing relationships "
        "with Meridian's top 150 enterprise accounts. Lisa Park joined in 2019 as a Software Engineer "
        "specializing in data pipeline architecture and was promoted to Senior Engineer in 2022."
    ))

    doc.add_page_break()

    add_heading_paragraph(doc, "C. Defendant's Breach and Misappropriation")

    add_body_paragraph(doc, (
        "On February 28, 2024, Dr. Ramaswamy submitted her resignation from Meridian, effective "
        "March 15, 2024. Mr. O'Brien resigned on March 1, 2024, and Ms. Park on March 4, 2024. "
        "All three cited 'personal reasons' for their departures. However, Meridian's investigation "
        "subsequently revealed that Cornerstone had extended offers of employment to all three "
        "individuals in January 2024—while they were still employed by Meridian."
    ))

    add_body_paragraph(doc, (
        "Forensic analysis of Meridian's internal systems, conducted by CyberGuard Analytics, Inc., "
        "revealed several alarming patterns of data exfiltration in the weeks preceding the "
        "departures. Specifically:"
    ))

    items = [
        ("Dr. Ramaswamy accessed and downloaded 4,217 files from the Project Phoenix repository "
         "between February 1 and February 27, 2024, including source code, algorithm specifications, "
         "and performance benchmarks. Her download volume during this period was approximately "
         "14 times her normal monthly average."),
        ("Mr. O'Brien exported client contact databases containing information on 2,347 corporate "
         "accounts, including contract values, renewal dates, pricing terms, and key decision-maker "
         "profiles. These exports occurred between February 15 and March 1, 2024."),
        ("Ms. Park accessed proprietary data pipeline schematics and infrastructure configuration "
         "files totaling 892 documents. She also forwarded 37 internal technical documents to a "
         "personal email address on February 22, 2024."),
    ]

    for item in items:
        add_body_paragraph(doc, f"  \u2022  {item}")

    add_body_paragraph(doc, (
        "On September 12, 2024, Cornerstone publicly announced the launch of DataForge Pro. "
        "Meridian's technical team conducted a detailed comparative analysis revealing that DataForge "
        "Pro's machine learning pipeline shares 89% structural similarity with Nexus's Project Phoenix "
        "architecture. The following specific parallels were identified by expert witness Dr. Alison "
        "Chen of Stanford University:"
    ))

    add_block_quote(doc, (
        "\"Having examined both the Nexus Analytics Suite and DataForge Pro, I conclude with a high "
        "degree of scientific certainty that DataForge Pro's predictive modeling engine is substantially "
        "derived from the Nexus platform. The algorithmic approach, including the specific implementation "
        "of gradient-boosted ensemble methods with custom regularization parameters, is virtually "
        "identical. The probability of independent parallel development producing this level of "
        "similarity is negligible.\" (Chen Decl. \u00b6 42.)"
    ))

    doc.add_page_break()

    # =============================================
    # III. LEGAL STANDARD (Pages 8-9)
    # =============================================
    add_heading_paragraph(doc, "III. LEGAL STANDARD FOR PRELIMINARY INJUNCTION")

    add_body_paragraph(doc, (
        "A plaintiff seeking a preliminary injunction in the Ninth Circuit must demonstrate: "
        "(1) a likelihood of success on the merits; (2) a likelihood of irreparable harm in the "
        "absence of preliminary relief; (3) that the balance of equities tips in the plaintiff's "
        "favor; and (4) that an injunction is in the public interest. Winter v. Natural Resources "
        "Defense Council, Inc., 555 U.S. 7, 20 (2008)."
    ))

    add_body_paragraph(doc, (
        "The Ninth Circuit employs a 'sliding scale' approach whereby a stronger showing on one "
        "element may offset a weaker showing on another, so long as all four elements are satisfied "
        "to some degree. Alliance for the Wild Rockies v. Cottrell, 632 F.3d 1127, 1131-35 (9th "
        "Cir. 2011). In trade secret cases, courts have recognized that the risk of continuing "
        "misappropriation weighs heavily in favor of injunctive relief."
    ))

    add_body_paragraph(doc, (
        "The Supreme Court has emphasized that injunctive relief is an extraordinary remedy that "
        "should not be granted as a matter of course. eBay Inc. v. MercExchange, LLC, 547 U.S. "
        "388, 391 (2006). Nevertheless, where a plaintiff demonstrates a clear entitlement to "
        "protection of its trade secrets, courts have consistently granted preliminary injunctions "
        "to prevent ongoing harm. See Cadence Design Systems, Inc. v. Avant! Corp., 125 F.3d 824, "
        "829 (9th Cir. 1997)."
    ))

    add_body_paragraph(doc, (
        "Under the Defend Trade Secrets Act ('DTSA'), 18 U.S.C. \u00a7 1836(b)(3)(A), a court "
        "may grant an injunction to prevent any actual or threatened misappropriation of trade "
        "secrets, provided the order does not 'prevent a person from entering into an employment "
        "relationship, and that conditions placed on such employment shall be based on evidence of "
        "threatened misappropriation and not merely on the information the person knows.'"
    ))

    doc.add_page_break()

    # =============================================
    # IV. ARGUMENT (Pages 9-14)
    # =============================================
    add_heading_paragraph(doc, "IV. ARGUMENT")

    add_heading_paragraph(doc, "A. Likelihood of Success on the Merits")

    add_body_paragraph(doc, (
        "Meridian is likely to prevail on its claims for trade secret misappropriation under both "
        "the DTSA and the California Uniform Trade Secrets Act ('CUTSA'), Cal. Civ. Code \u00a7 3426 "
        "et seq. To establish a trade secret misappropriation claim, a plaintiff must show: "
        "(1) the existence of a trade secret; (2) the defendant's misappropriation of the trade "
        "secret; and (3) resulting or threatened damage to the plaintiff."
    ))

    add_heading_paragraph(doc, "1. Existence of Trade Secrets")

    add_body_paragraph(doc, (
        "Meridian's Nexus Analytics Suite constitutes protectable trade secrets under both federal "
        "and state law. The DTSA defines a trade secret as 'all forms and types of financial, "
        "business, scientific, technical, economic, or engineering information' that the owner has "
        "taken 'reasonable measures to keep secret' and that derives 'independent economic value "
        "from not being generally known.' 18 U.S.C. \u00a7 1839(3)."
    ))

    add_body_paragraph(doc, (
        "Meridian has implemented extensive measures to protect the confidentiality of the Nexus "
        "platform, including: (a) requiring all employees with access to sign non-disclosure "
        "agreements; (b) implementing role-based access controls limiting access to source code "
        "repositories; (c) maintaining encrypted development environments with multi-factor "
        "authentication; (d) conducting regular security audits by independent cybersecurity firms; "
        "and (e) marking all proprietary documents as 'Confidential' or 'Highly Confidential.'"
    ))

    add_body_paragraph(doc, (
        "The independent economic value of Meridian's trade secrets is substantial. The Nexus "
        "platform's predictive analytics capabilities provide Meridian with a significant competitive "
        "advantage, enabling it to achieve prediction accuracy rates 23% higher than industry "
        "benchmarks. This technological edge has been a primary factor in Meridian's ability to "
        "command premium pricing and achieve a customer retention rate of 94.7%."
    ))

    add_heading_paragraph(doc, "2. Misappropriation")

    add_body_paragraph(doc, (
        "The evidence of misappropriation is compelling. The departing employees had access to "
        "Meridian's most sensitive trade secrets by virtue of their roles. The forensic evidence "
        "demonstrates extensive downloading and exfiltration of proprietary data in the weeks before "
        "their departures. The Ninth Circuit has held that such evidence of mass data downloads "
        "prior to departure creates a strong inference of misappropriation."
    ))

    add_block_quote(doc, (
        "\"Where former employees download large volumes of confidential data shortly before joining "
        "a competitor, and that competitor subsequently launches a substantially similar product, "
        "the inference of misappropriation is not merely reasonable—it is compelling.\" MAI Systems "
        "Corp. v. Peak Computer, Inc., 991 F.2d 511, 521 (9th Cir. 1993)."
    ))

    add_body_paragraph(doc, (
        "Moreover, the structural similarity between DataForge Pro and the Nexus platform, as "
        "documented by Dr. Chen's expert analysis, provides powerful circumstantial evidence that "
        "Cornerstone used Meridian's trade secrets in developing its competing product. The 89% "
        "structural similarity identified by Dr. Chen is far beyond what could be attributed to "
        "coincidence or independent development."
    ))

    doc.add_page_break()

    add_heading_paragraph(doc, "B. Irreparable Harm")

    add_body_paragraph(doc, (
        "Meridian will suffer irreparable harm absent injunctive relief. In trade secret cases, "
        "courts in this circuit have consistently recognized that the ongoing use and disclosure of "
        "misappropriated trade secrets constitutes irreparable harm per se. See Ruckelshaus v. "
        "Monsanto Co., 467 U.S. 986, 1012 (1984) ('The right to exclude others from using trade "
        "secrets is central to the very definition of the property right')."
    ))

    add_body_paragraph(doc, (
        "The harm to Meridian is irreparable for several independent reasons. First, trade secrets, "
        "once disclosed, lose their protected status permanently. Unlike patent infringement, which "
        "can be remedied by monetary damages and prospective injunctions, the dissemination of trade "
        "secrets cannot be undone. Second, Meridian faces imminent loss of market share and customer "
        "relationships that would be difficult or impossible to quantify with precision."
    ))

    add_body_paragraph(doc, (
        "Since the launch of DataForge Pro, Meridian has lost seven major enterprise contracts "
        "totaling approximately $18.3 million in annual recurring revenue. Three additional clients "
        "representing $7.2 million in annual revenue have indicated they are evaluating DataForge "
        "Pro as an alternative. The erosion of Meridian's market position, cultivated over a decade "
        "of investment and innovation, cannot be adequately compensated through monetary damages alone."
    ))

    add_body_paragraph(doc, (
        "Third, allowing Cornerstone to continue using Meridian's trade secrets would enable "
        "Cornerstone to further develop and improve upon the misappropriated technology, compounding "
        "the competitive harm to Meridian. Each day that passes without injunctive relief allows "
        "Cornerstone to deepen its understanding and exploitation of Meridian's proprietary "
        "innovations, making eventual remediation increasingly difficult."
    ))

    add_heading_paragraph(doc, "C. Balance of Equities")

    add_body_paragraph(doc, (
        "The balance of equities tips decidedly in Meridian's favor. On one side of the scale "
        "stands Meridian, a company that has invested over $147 million and nearly a decade of "
        "sustained effort in developing its proprietary technology. On the other stands Cornerstone, "
        "a company that stands accused of building its competing product on a foundation of stolen "
        "trade secrets."
    ))

    add_body_paragraph(doc, (
        "Any hardship to Cornerstone from an injunction is self-inflicted. A party that achieves "
        "its market position through misappropriation cannot claim equitable hardship when asked to "
        "cease using the misappropriated materials. See DVD Copy Control Ass'n v. Bunner, 116 Cal. "
        "App. 4th 241, 254 (2004) ('One who misappropriates a trade secret cannot be heard to "
        "complain of the hardship caused by an injunction against further use')."
    ))

    add_body_paragraph(doc, (
        "Furthermore, the scope of the requested injunction is narrowly tailored. Meridian does not "
        "seek to prevent Cornerstone from operating in the enterprise analytics market entirely. "
        "Rather, Meridian requests only that Cornerstone be enjoined from using, disclosing, or "
        "benefiting from Meridian's specific trade secrets, and that the departing employees be "
        "required to comply with their contractual obligations."
    ))

    doc.add_page_break()

    add_heading_paragraph(doc, "D. Public Interest")

    add_body_paragraph(doc, (
        "The public interest strongly favors granting the requested preliminary injunction. The "
        "protection of trade secrets serves vital public policy objectives. Trade secret protection "
        "incentivizes innovation by ensuring that companies can reap the benefits of their "
        "investments in research and development. Without such protection, the incentive to invest "
        "in cutting-edge technology would be severely diminished."
    ))

    add_body_paragraph(doc, (
        "Congress recognized the importance of trade secret protection when it enacted the Defend "
        "Trade Secrets Act in 2016, creating a federal civil cause of action for trade secret "
        "misappropriation. The legislative history of the DTSA emphasizes that 'trade secrets are "
        "an integral part of a company's competitive advantage' and that their protection 'promotes "
        "innovation and economic growth.' S. Rep. No. 114-220, at 2 (2016)."
    ))

    add_body_paragraph(doc, (
        "Moreover, the enforcement of contractual obligations—including non-disclosure agreements "
        "and invention assignment clauses—serves the public interest by promoting commercial "
        "certainty and the sanctity of contracts. Allowing employees to disregard their contractual "
        "obligations with impunity would undermine the foundation of employment relationships in "
        "technology industries."
    ))

    add_body_paragraph(doc, (
        "The public interest is also served by maintaining fair competition. When a competitor gains "
        "market position through misappropriation rather than legitimate innovation, it distorts the "
        "competitive landscape and harms consumers who are deprived of the benefits of genuine "
        "competition. An injunction here would restore the level playing field that the law demands."
    ))

    doc.add_page_break()

    # =============================================
    # V. CONCLUSION (Page 14-15)
    # =============================================
    add_heading_paragraph(doc, "V. CONCLUSION")

    add_body_paragraph(doc, (
        "For the foregoing reasons, Plaintiff Meridian Technologies, Inc. respectfully requests "
        "that this Court grant its motion for a preliminary injunction and enter an order:"
    ))

    orders = [
        "Enjoining Cornerstone Data Solutions, LLC, its officers, agents, employees, and all persons acting in concert with them, from using, disclosing, or benefiting from any of Meridian's trade secrets, including but not limited to the Nexus Analytics Suite source code, algorithms, customer databases, and pricing models;",
        "Requiring Cornerstone to identify and sequester all materials derived from or incorporating Meridian's trade secrets, subject to inspection by a court-appointed special master;",
        "Enjoining Dr. Priya Ramaswamy, Kevin O'Brien, and Lisa Park from disclosing any of Meridian's confidential information to Cornerstone or any third party, and from performing work for Cornerstone that draws upon or relates to knowledge gained during their employment with Meridian;",
        "Requiring the return or certified destruction of all copies of Meridian's proprietary materials in the possession, custody, or control of Cornerstone and the individual defendants;",
        "Awarding Meridian its costs and attorneys' fees incurred in bringing this motion; and",
        "Granting such other and further relief as this Court deems just and proper.",
    ]

    for i, order in enumerate(orders, 1):
        add_body_paragraph(doc, f"  ({i})  {order}")

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "")

    p = add_body_paragraph(doc, "Dated: March 15, 2025")
    add_body_paragraph(doc, "")

    p = add_body_paragraph(doc, "Respectfully submitted,")
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "")

    p = add_body_paragraph(doc, "________________________________")
    p = add_body_paragraph(doc, "CATHERINE A. MORRISON, ESQ.")
    p = add_body_paragraph(doc, "DAVID R. NAKAMURA, ESQ.")
    p = add_body_paragraph(doc, "Morrison & Nakamura LLP")
    p = add_body_paragraph(doc, "One Market Plaza, Suite 3200")
    p = add_body_paragraph(doc, "San Francisco, California 94105")
    p = add_body_paragraph(doc, "Tel: (415) 555-7890")
    p = add_body_paragraph(doc, "Fax: (415) 555-7891")
    p = add_body_paragraph(doc, "Email: cmorrison@morrisonak.com")
    add_body_paragraph(doc, "")
    p = add_body_paragraph(doc, "Attorneys for Plaintiff")
    p = add_body_paragraph(doc, "MERIDIAN TECHNOLOGIES, INC.")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
