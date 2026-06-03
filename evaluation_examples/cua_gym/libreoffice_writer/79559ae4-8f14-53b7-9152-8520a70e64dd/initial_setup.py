"""
Initial Setup: Sort table alphabetically by 'Last Name' column
Task ID: writer_tbl_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'  # VM path — file placed on Desktop per task context
TASK_ID = 'contacts_sort'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a title paragraph for context
    title = doc.add_paragraph("Contact Directory")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Add the contacts table (6 rows x 3 columns)
    # Row 1 is header; rows 2-6 are data in UNSORTED order
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    # Row 1: Header
    header_row = table.rows[0]
    header_row.cells[0].text = 'Last Name'
    header_row.cells[1].text = 'First Name'
    header_row.cells[2].text = 'Phone'
    # Bold the header cells
    for cell in header_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows in UNSORTED (original) order — task requires sorting them
    data = [
        ('Zhang',    'Wei',     '555-0901'),
        ('Garcia',   'Maria',   '555-0234'),
        ('Anderson', 'James',   '555-0567'),
        ('Wilson',   'Sarah',   '555-0789'),
        ('Brown',    'Michael', '555-0456'),
    ]

    for i, (last, first, phone) in enumerate(data, start=1):
        row = table.rows[i]
        row.cells[0].text = last
        row.cells[1].text = first
        row.cells[2].text = phone

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
