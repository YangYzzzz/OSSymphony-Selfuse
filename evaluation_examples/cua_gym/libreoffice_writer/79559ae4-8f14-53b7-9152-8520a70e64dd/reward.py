"""
Reward Script: Sort the table alphabetically by the 'Last Name' column in ascending order (A to Z).
Task ID: writer_tbl_031
Domain: libreoffice_writer
Scoring:
  Component 1: Last Name column is sorted A to Z (0.5 pts)
  Component 2: All 5 data rows match expected sorted order with correct First Name and Phone pairings (0.5 pts)
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_031'

# Expected sorted data rows (header excluded)
EXPECTED_ROWS = [
    ['Anderson', 'James', '555-0567'],
    ['Brown', 'Michael', '555-0456'],
    ['Garcia', 'Maria', '555-0234'],
    ['Wilson', 'Sarah', '555-0789'],
    ['Zhang', 'Wei', '555-0901'],
]

def verify_task(file_path):
    """
    Verify that the contacts_sort.docx table is sorted alphabetically by Last Name (A to Z).
    Returns a float between 0.0 and 1.0.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: table must exist with expected structure
    if not doc.tables:
        print("CRITICAL: No tables found in document.")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    rows = table.rows

    if len(rows) < 2:
        print("CRITICAL: Table has fewer than 2 rows.")
        print("REWARD: 0.0")
        return 0.0

    # Get data rows (skip header row 0)
    data_rows = []
    for i in range(1, len(rows)):
        row_data = [cell.text.strip() for cell in rows[i].cells]
        # Trim to first 3 cols if more
        row_data = row_data[:3] if len(row_data) >= 3 else row_data
        data_rows.append(row_data)

    # Component 1: Last Name column is sorted A-Z (0.5 points)
    # This checks whether the first column of data rows is in ascending alphabetical order.
    # On initial_env, order is: Zhang, Garcia, Anderson, Wilson, Brown (NOT sorted) → FAIL
    # On golden_env, order is: Anderson, Brown, Garcia, Wilson, Zhang (sorted A-Z) → PASS
    try:
        last_names = [row[0] for row in data_rows if len(row) >= 1]
        expected_last_names = [r[0] for r in EXPECTED_ROWS]
        if last_names == expected_last_names:
            print(f"PASS: Component 1 — Last Name column sorted A-Z: {last_names} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Last Name column not sorted A-Z. Found: {last_names}, Expected: {expected_last_names}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 5 data rows match expected sorted rows exactly (correct pairings) (0.5 points)
    # This checks that each row's First Name and Phone are still paired with the correct Last Name.
    # On initial_env: pairings exist but in wrong order → FAIL (rows don't match expected sorted rows)
    # On golden_env: rows match expected sorted rows exactly → PASS
    try:
        if len(data_rows) == len(EXPECTED_ROWS):
            all_match = all(
                data_rows[i] == EXPECTED_ROWS[i]
                for i in range(len(EXPECTED_ROWS))
            )
            if all_match:
                print(f"PASS: Component 2 — All 5 data rows match expected sorted order with correct pairings (0.5 pts)")
                total_score += 0.5
            else:
                mismatches = [
                    f"  Row {i+1}: got {data_rows[i]}, expected {EXPECTED_ROWS[i]}"
                    for i in range(len(EXPECTED_ROWS))
                    if data_rows[i] != EXPECTED_ROWS[i]
                ]
                print(f"FAIL: Component 2 — Row pairings don't match expected. Mismatches:")
                for m in mismatches:
                    print(m)
        else:
            print(f"FAIL: Component 2 — Expected {len(EXPECTED_ROWS)} data rows, found {len(data_rows)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


file_path = f'{WORKDIR}/contacts_sort.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
