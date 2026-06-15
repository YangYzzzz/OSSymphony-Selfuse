"""
Initial Setup: Create client_feedback.docx and mockup.png on the Desktop
Task ID: osworld_multi_apps_writer_to_gimp_008
Domain: libreoffice_writer + gimp (multi-app)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_to_gimp_008'

DOCX_PATH = f'{WORKDIR}/client_feedback.docx'
MOCKUP_PATH = f'{WORKDIR}/mockup.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_docx():
    """Create client_feedback.docx with specific visual change requests."""
    doc = Document()

    # Title
    title = doc.add_heading("Client Feedback — Product Mockup Revision", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Dear Design Team,"
    )
    intro_run.font.name = "Calibri"
    intro_run.font.size = Pt(11)

    doc.add_paragraph()

    body = doc.add_paragraph()
    body_run = body.add_run(
        "After reviewing the latest mockup, our stakeholders have requested the following specific changes "
        "to 'mockup.png'. Please apply all of the modifications listed below and save the revised image "
        "as 'mockup_revised.png' on the Desktop."
    )
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(11)

    doc.add_paragraph()

    # Section heading
    changes_heading = doc.add_heading("Requested Modifications:", level=2)

    # Modification 1: Change background to white
    item1 = doc.add_paragraph(style="List Number")
    run1a = item1.add_run("Change background color to white")
    run1a.bold = True
    run1a.font.name = "Calibri"
    run1a.font.size = Pt(11)
    run1b = item1.add_run(
        " — The current dark/colored background should be replaced with a clean white (#FFFFFF) background. "
        "All foreground elements (shapes, text) should be retained."
    )
    run1b.font.name = "Calibri"
    run1b.font.size = Pt(11)

    # Modification 2: Scale image by 50%
    item2 = doc.add_paragraph(style="List Number")
    run2a = item2.add_run("Scale the image to 50% of its original dimensions")
    run2a.bold = True
    run2a.font.name = "Calibri"
    run2a.font.size = Pt(11)
    run2b = item2.add_run(
        " — Resize the canvas and all content to exactly 50% of the current width and height. "
        "Maintain the aspect ratio during resizing."
    )
    run2b.font.name = "Calibri"
    run2b.font.size = Pt(11)

    # Modification 3: Apply blur
    item3 = doc.add_paragraph(style="List Number")
    run3a = item3.add_run("Apply a Gaussian blur")
    run3a.bold = True
    run3a.font.name = "Calibri"
    run3a.font.size = Pt(11)
    run3b = item3.add_run(
        " — Apply a Gaussian blur with a radius of 2 pixels to soften the overall image. "
        "This will give the mockup a smoother, polished appearance."
    )
    run3b.font.name = "Calibri"
    run3b.font.size = Pt(11)

    doc.add_paragraph()

    # Closing
    closing = doc.add_paragraph()
    closing_run = closing.add_run(
        "Please ensure all three modifications are applied to the final file. "
        "The output file must be saved as 'mockup_revised.png' on the Desktop."
    )
    closing_run.font.name = "Calibri"
    closing_run.font.size = Pt(11)

    doc.add_paragraph()

    sign = doc.add_paragraph()
    sign_run = sign.add_run("Best regards,\nAcme Product Design Team")
    sign_run.font.name = "Calibri"
    sign_run.font.size = Pt(11)

    doc.save(DOCX_PATH)
    print(f"Created: {DOCX_PATH}")


def create_mockup():
    """Create mockup.png — a realistic product mockup image with colored background and shapes."""
    # Image size: 800x600
    width, height = 800, 600

    # Background: dark navy blue
    bg_color = (30, 45, 90)
    img = Image.new("RGB", (width, height), color=bg_color)
    draw = ImageDraw.Draw(img)

    # --- Product card area (white rounded-rect simulation) ---
    card_left, card_top, card_right, card_bottom = 100, 80, 700, 520
    draw.rectangle([card_left, card_top, card_right, card_bottom],
                   fill=(255, 255, 255), outline=(200, 200, 210), width=2)

    # --- Header bar on card ---
    draw.rectangle([card_left, card_top, card_right, card_top + 60],
                   fill=(52, 120, 220))

    # Try to load a font; fall back to default if unavailable
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
        font_label = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
    except (IOError, OSError):
        font_title = ImageFont.load_default()
        font_label = font_title
        font_small = font_title

    # Title text in header bar
    draw.text((120, 90), "Product Mockup v2.3", fill=(255, 255, 255), font=font_title)

    # --- Product image placeholder (colored rectangle) ---
    draw.rectangle([120, 160, 420, 400], fill=(255, 160, 30), outline=(200, 120, 10), width=3)
    draw.text((190, 265), "Product Image", fill=(80, 40, 0), font=font_label)

    # --- Product details panel ---
    draw.text((450, 165), "Model X-500", fill=(30, 30, 80), font=font_label)
    draw.text((450, 200), "Category: Electronics", fill=(60, 60, 100), font=font_small)
    draw.text((450, 230), "SKU: EL-X500-BLK", fill=(60, 60, 100), font=font_small)
    draw.text((450, 260), "Price: $299.99", fill=(180, 30, 30), font=font_label)
    draw.text((450, 295), "Stock: 142 units", fill=(30, 120, 30), font=font_small)
    draw.text((450, 325), "Rating: 4.7 / 5.0", fill=(60, 60, 100), font=font_small)

    # Divider line
    draw.line([(440, 360), (690, 360)], fill=(200, 200, 215), width=2)

    # CTA button simulation
    draw.rectangle([450, 375, 680, 415], fill=(52, 120, 220), outline=(30, 90, 180), width=2)
    draw.text((510, 384), "Add to Cart", fill=(255, 255, 255), font=font_label)

    # --- Bottom info bar ---
    draw.rectangle([card_left, card_bottom - 40, card_right, card_bottom],
                   fill=(240, 240, 248))
    draw.text((120, card_bottom - 30), "© 2025 Acme Corp  |  All rights reserved",
               fill=(100, 100, 120), font=font_small)

    # --- Decorative circles in background ---
    draw.ellipse([10, 10, 80, 80], fill=(50, 70, 130), outline=(70, 90, 160))
    draw.ellipse([720, 520, 790, 590], fill=(50, 70, 130), outline=(70, 90, 160))

    img.save(MOCKUP_PATH)
    print(f"Created: {MOCKUP_PATH}")


def main():
    os.makedirs(WORKDIR, exist_ok=True)
    create_docx()
    create_mockup()

    # GUI-ready startup: open the docx in LibreOffice Writer and mockup.png in GIMP
    launch_gui(f'libreoffice --writer "{DOCX_PATH}"', delay_sec=2.0)
    launch_gui(f'gimp "{MOCKUP_PATH}"', delay_sec=2.0)

    print('GUI_READY: launched LibreOffice Writer and GIMP with DISPLAY=:0')


main()
