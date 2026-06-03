"""
Reward Script: Apply heading styles to create a complete document outline
Task ID: writer_struct_065
Domain: libreoffice_writer
Scoring:
  Component 1: 'Chapter 1: Introduction' and 'Chapter 2: Analysis' styled as Heading 1  (0.30 pts)
  Component 2: 'Background', 'Scope', 'Financial Analysis', 'Market Analysis' styled as Heading 2 (0.40 pts)
  Component 3: 'Revenue Trends' styled as Heading 3  (0.20 pts)
  Component 4: Body text paragraphs remain as Normal style (no over-heading)  (0.10 pts)
  Total: 1.00
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_065'

# Expected heading assignments
HEADING1_TITLES = ['Chapter 1: Introduction', 'Chapter 2: Analysis']
HEADING2_TITLES = ['Background', 'Scope', 'Financial Analysis', 'Market Analysis']
HEADING3_TITLES = ['Revenue Trends']

# Body paragraphs that must remain Normal (not reclassified as headings)
BODY_TEXT_SNIPPETS = [
    'This report provides a comprehensive examination',
    'The analysis draws upon internal financial records',
    'Meridian Technologies Inc. was founded in 2011',
    'Revenue has grown at a compound annual rate',
    'Key milestones include the acquisition of DataFlow',
    'This analysis covers the following areas',
    'Geographic coverage includes operations in Singapore',
    'All monetary values are expressed in United States',
    'The following sections present a detailed examination',
    'Total revenue for the nine-month period ended 30 September',
    'Gross profit margin improved by 2.3 percentage points',
    'Net income for the period was USD 18.6 million',
    'Monthly recurring revenue (MRR) has demonstrated',
    'Customer concentration risk has diminished',
    'Annualised contract value (ACV) of new business signed',
    'The Asia-Pacific enterprise software market is projected',
    'Competitive intensity remains high',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build a lookup: paragraph text -> style name
    para_style = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            para_style[text] = para.style.name

    # Component 1: Heading 1 — 'Chapter 1: Introduction' and 'Chapter 2: Analysis' (0.30 points)
    try:
        h1_correct = []
        h1_wrong = []
        for title in HEADING1_TITLES:
            # Find the paragraph by exact text match
            found_style = para_style.get(title)
            if found_style == 'Heading 1':
                h1_correct.append(title)
            else:
                h1_wrong.append(f"'{title}' has style '{found_style}' (expected 'Heading 1')")

        h1_pts = round(0.30 * len(h1_correct) / len(HEADING1_TITLES), 2)
        if len(h1_correct) == len(HEADING1_TITLES):
            print(f"PASS: Component 1 — Both Heading 1 titles correctly styled: {h1_correct} (0.30 pts)")
            total_score += h1_pts
        elif len(h1_correct) > 0:
            print(f"PARTIAL: Component 1 — {len(h1_correct)}/{len(HEADING1_TITLES)} Heading 1 correct "
                  f"(+{h1_pts} pts). Issues: {h1_wrong}")
            total_score += h1_pts
        else:
            print(f"FAIL: Component 1 — No Heading 1 titles found. Issues: {h1_wrong}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Heading 2 — 'Background', 'Scope', 'Financial Analysis', 'Market Analysis' (0.40 points)
    try:
        h2_correct = []
        h2_wrong = []
        for title in HEADING2_TITLES:
            found_style = para_style.get(title)
            if found_style == 'Heading 2':
                h2_correct.append(title)
            else:
                h2_wrong.append(f"'{title}' has style '{found_style}' (expected 'Heading 2')")

        h2_pts = round(0.40 * len(h2_correct) / len(HEADING2_TITLES), 2)
        if len(h2_correct) == len(HEADING2_TITLES):
            print(f"PASS: Component 2 — All 4 Heading 2 titles correctly styled: {h2_correct} (0.40 pts)")
            total_score += h2_pts
        elif len(h2_correct) > 0:
            print(f"PARTIAL: Component 2 — {len(h2_correct)}/{len(HEADING2_TITLES)} Heading 2 correct "
                  f"(+{h2_pts} pts). Issues: {h2_wrong}")
            total_score += h2_pts
        else:
            print(f"FAIL: Component 2 — No Heading 2 titles found. Issues: {h2_wrong}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Heading 3 — 'Revenue Trends' (0.20 points)
    try:
        found_style = para_style.get('Revenue Trends')
        if found_style == 'Heading 3':
            print(f"PASS: Component 3 — 'Revenue Trends' correctly styled as Heading 3 (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — 'Revenue Trends' has style '{found_style}' (expected 'Heading 3')")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Exactly 7 headings total (2 H1 + 4 H2 + 1 H3) with correct titles and no over-application (0.10 points)
    # This verifies that ONLY the 7 specified titles got heading styles, and NO body text was erroneously promoted.
    # In initial_env, total headings = 0 (not 7), so this correctly fails.
    try:
        expected_headings = set(HEADING1_TITLES + HEADING2_TITLES + HEADING3_TITLES)
        all_heading_paras = [(p.text.strip(), p.style.name) for p in doc.paragraphs
                             if p.style.name.startswith('Heading') and p.text.strip()]
        total_headings = len(all_heading_paras)
        unexpected = [(t, s) for t, s in all_heading_paras if t not in expected_headings]

        if total_headings == 7 and len(unexpected) == 0:
            print(f"PASS: Component 4 — Exactly 7 heading paragraphs, all matching expected titles (0.10 pts)")
            total_score += 0.10
        elif total_headings == 0:
            print(f"FAIL: Component 4 — No heading paragraphs found (expected 7)")
        elif unexpected:
            print(f"FAIL: Component 4 — Unexpected heading styles on body text: {unexpected}")
        else:
            print(f"FAIL: Component 4 — Expected 7 heading paragraphs, found {total_headings}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in a given env
file_path = '/home/user/Desktop/business_analysis.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
