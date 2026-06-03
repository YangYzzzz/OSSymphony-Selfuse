"""
Initial Setup: Apply heading styles to create a complete document outline in a business analysis report.
Task ID: writer_struct_065
Domain: libreoffice_writer

Creates a 12-section business analysis document where all 7 section titles
are styled as 'Default Paragraph Style' (no heading styles) — the agent
must apply Heading 1/2/3 to them.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'business_analysis'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_paragraph(doc, text, indent=False):
    """Add a normal body paragraph using Default Paragraph Style."""
    para = doc.add_paragraph(text)
    para.style = doc.styles['Normal']
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    return para


def add_section_title(doc, text):
    """Add a section title styled as Normal (Default Paragraph Style) — NOT a heading."""
    para = doc.add_paragraph(text)
    para.style = doc.styles['Normal']
    # Bold to make it visually look like a title without using heading styles
    for run in para.runs:
        run.bold = True
    return para


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Document Title (using Normal style, not a heading) ---
    title_para = doc.add_paragraph('Business Analysis Report 2025')
    title_para.style = doc.styles['Normal']
    title_run = title_para.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph('')  # spacer

    # =========================================================
    # CHAPTER 1: INTRODUCTION  (currently Normal/Default style)
    # =========================================================
    ch1_para = doc.add_paragraph('Chapter 1: Introduction')
    ch1_para.style = doc.styles['Normal']
    for run in ch1_para.runs:
        run.bold = True
        run.font.size = Pt(14)

    add_body_paragraph(doc,
        'This report provides a comprehensive examination of the current business '
        'environment, strategic positioning, and financial performance of Meridian '
        'Technologies Inc. for the fiscal year 2024-2025. The findings presented '
        'herein are based on data collected from Q1 through Q3 of the reporting period.')
    add_body_paragraph(doc,
        'The analysis draws upon internal financial records, market intelligence '
        'gathered by the Strategy & Research division, and benchmarking data from '
        'comparable industry participants.')

    doc.add_paragraph('')  # spacer

    # --- Background (currently Normal/Default style) ---
    bg_para = doc.add_paragraph('Background')
    bg_para.style = doc.styles['Normal']
    for run in bg_para.runs:
        run.bold = True
        run.font.size = Pt(13)

    add_body_paragraph(doc,
        'Meridian Technologies Inc. was founded in 2011 with the mission to '
        'deliver enterprise-grade software solutions to mid-market companies across '
        'the Asia-Pacific region. Over the past fourteen years the company has grown '
        'from a three-person startup to an organisation of approximately 1,200 '
        'full-time employees operating across seven countries.')
    add_body_paragraph(doc,
        'Revenue has grown at a compound annual rate of 18.4% since 2018, driven '
        'primarily by expansion in cloud subscription services and professional '
        'consulting engagements. The company completed its Series D funding round '
        'in November 2023, raising USD 95 million to accelerate product development '
        'and geographic expansion.')
    add_body_paragraph(doc,
        'Key milestones include the acquisition of DataFlow Analytics in 2021, '
        'the launch of the MeridianOS platform in 2022, and a strategic partnership '
        'with NovaBridge Capital announced in February 2024.')

    doc.add_paragraph('')  # spacer

    # --- Scope (currently Normal/Default style) ---
    scope_para = doc.add_paragraph('Scope')
    scope_para.style = doc.styles['Normal']
    for run in scope_para.runs:
        run.bold = True
        run.font.size = Pt(13)

    add_body_paragraph(doc,
        'This analysis covers the following areas: financial performance metrics '
        'including revenue, gross margin, EBITDA, and cash flow; market dynamics '
        'and competitive landscape across Meridian\'s three primary verticals; '
        'operational efficiency indicators; and strategic initiatives planned for '
        'the remainder of FY2025.')
    add_body_paragraph(doc,
        'Geographic coverage includes operations in Singapore, Australia, Japan, '
        'South Korea, India, Vietnam, and the Philippines. The scope excludes '
        'activities related to the recently announced entry into the Middle East '
        'market, which remains under regulatory review.')
    add_body_paragraph(doc,
        'All monetary values are expressed in United States Dollars (USD) unless '
        'otherwise indicated. Year-over-year comparisons reference FY2023-2024 as '
        'the baseline period.')

    doc.add_paragraph('')  # spacer

    # =========================================================
    # CHAPTER 2: ANALYSIS  (currently Normal/Default style)
    # =========================================================
    ch2_para = doc.add_paragraph('Chapter 2: Analysis')
    ch2_para.style = doc.styles['Normal']
    for run in ch2_para.runs:
        run.bold = True
        run.font.size = Pt(14)

    add_body_paragraph(doc,
        'The following sections present a detailed examination of Meridian\'s '
        'financial position and market standing. Each subsection is supported by '
        'quantitative data and qualitative commentary from senior management.')

    doc.add_paragraph('')  # spacer

    # --- Financial Analysis (currently Normal/Default style) ---
    fin_para = doc.add_paragraph('Financial Analysis')
    fin_para.style = doc.styles['Normal']
    for run in fin_para.runs:
        run.bold = True
        run.font.size = Pt(13)

    add_body_paragraph(doc,
        'Total revenue for the nine-month period ended 30 September 2025 was '
        'USD 187.3 million, representing a 22.1% increase compared with the '
        'equivalent period in the prior fiscal year. Cloud subscription revenue '
        'constituted 68% of total revenue at USD 127.4 million, while professional '
        'services contributed the remaining 32% at USD 59.9 million.')
    add_body_paragraph(doc,
        'Gross profit margin improved by 2.3 percentage points to 61.8%, reflecting '
        'favourable product mix shift toward higher-margin subscription offerings '
        'and improved delivery efficiency in the services segment. Operating expenses '
        'increased by 14.7% year-on-year, primarily due to headcount additions in '
        'the sales and engineering functions.')
    add_body_paragraph(doc,
        'Net income for the period was USD 18.6 million, compared with USD 11.2 million '
        'in the prior year period. Earnings per share on a diluted basis were USD 0.42, '
        'up from USD 0.26. Cash and cash equivalents at period-end stood at '
        'USD 143.7 million.')

    doc.add_paragraph('')  # spacer

    # --- Revenue Trends (currently Normal/Default style) ---
    rev_para = doc.add_paragraph('Revenue Trends')
    rev_para.style = doc.styles['Normal']
    for run in rev_para.runs:
        run.bold = True
        run.font.size = Pt(12)

    add_body_paragraph(doc,
        'Monthly recurring revenue (MRR) has demonstrated consistent upward '
        'trajectory over the reporting period. MRR reached USD 14.2 million in '
        'September 2025, up from USD 10.8 million in October 2024, representing '
        'a 31.5% increase over the eleven-month span.')
    add_body_paragraph(doc,
        'Customer concentration risk has diminished as the top-ten customers now '
        'account for 28.4% of total revenue, down from 34.1% in the prior year. '
        'Net revenue retention rate stood at 118% for the trailing twelve months, '
        'indicating significant upsell and expansion activity within the existing '
        'customer base.')
    add_body_paragraph(doc,
        'Annualised contract value (ACV) of new business signed in Q3 2025 was '
        'USD 22.7 million, the highest quarterly figure recorded by the company. '
        'Average deal size has increased by 17% year-on-year, partly attributable '
        'to the introduction of enterprise tier pricing in January 2025.')

    doc.add_paragraph('')  # spacer

    # --- Market Analysis (currently Normal/Default style) ---
    mkt_para = doc.add_paragraph('Market Analysis')
    mkt_para.style = doc.styles['Normal']
    for run in mkt_para.runs:
        run.bold = True
        run.font.size = Pt(13)

    add_body_paragraph(doc,
        'The Asia-Pacific enterprise software market is projected to reach '
        'USD 312 billion by 2028, expanding at a CAGR of 11.3% from 2023. '
        'Cloud-native solutions are displacing legacy on-premises systems at an '
        'accelerating pace, creating significant runway for providers with proven '
        'cloud offerings such as Meridian.')
    add_body_paragraph(doc,
        'Competitive intensity remains high. Primary competitors include regional '
        'challengers AxioSync Solutions and PrismPath Technologies, as well as '
        'global incumbents Oracle and SAP who continue to invest in cloud migration '
        'pathways for their installed customer bases.')
    add_body_paragraph(doc,
        'Meridian\'s net promoter score of 62 as of Q3 2025 compares favourably '
        'against the industry median of 44, suggesting strong customer satisfaction '
        'and advocacy levels. Brand recognition surveys conducted in six of the seven '
        'operating markets show aided awareness above 70% among IT decision-makers.')
    add_body_paragraph(doc,
        'Regulatory developments, including forthcoming data residency requirements '
        'in Japan and South Korea, may necessitate incremental infrastructure '
        'investment in 2026. Management has begun scoping work on in-country data '
        'centre deployments to ensure compliance ahead of the anticipated regulatory '
        'effective dates.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
