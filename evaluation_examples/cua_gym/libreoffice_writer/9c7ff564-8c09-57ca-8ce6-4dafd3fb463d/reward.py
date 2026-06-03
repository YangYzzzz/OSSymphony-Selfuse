"""
Reward Script: Generate a table of contents for the entire master document
Task ID: writer_rm_057
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): TOC field code exists in the document
  Component 2 (0.25): TOC field covers heading levels 1 through 3
  Component 3 (0.20): A TOC heading/title paragraph exists near the top of the document
  Component 4 (0.20): TOC is positioned before the main body content
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_057'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires inserting a Table of Contents that includes
    headings up to level 3. We verify:
    1. A TOC field code exists in the document XML
    2. The TOC field covers levels 1-3
    3. A TOC heading/title paragraph appears near the top
    4. The TOC is positioned before the main body content
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Parse XML namespace
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body

    # Component 1: TOC field code exists in the document (0.35 points)
    # A real TOC in .docx is represented by instrText containing "TOC"
    try:
        fld_instructions = body.findall('.//w:instrText', ns)
        toc_fields = [fi for fi in fld_instructions if fi.text and 'TOC' in fi.text.upper()]
        if len(toc_fields) > 0:
            toc_instr = toc_fields[0].text
            print(f"PASS: Component 1 — TOC field code found: {toc_instr!r} (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — No TOC field code found in document XML. "
                  f"Found {len(fld_instructions)} instrText fields, none containing 'TOC'.")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: TOC field covers heading levels 1 through 3 (0.25 points)
    # The standard TOC field for levels 1-3 includes \o "1-3"
    try:
        if len(toc_fields) > 0:
            toc_instr = toc_fields[0].text
            # Check for outline levels covering at least 1-3
            # Common patterns: \o "1-3", \o "1-4", etc.
            level_match = re.search(r'\\o\s*"(\d)-(\d)"', toc_instr)
            if level_match:
                start_level = int(level_match.group(1))
                end_level = int(level_match.group(2))
                if start_level <= 1 and end_level >= 3:
                    print(f"PASS: Component 2 — TOC covers levels {start_level}-{end_level}, "
                          f"includes all required levels 1-3 (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 2 — TOC covers levels {start_level}-{end_level}, "
                          f"does not fully include 1-3")
            else:
                # Some TOC fields use heading styles without \o switch
                # Check for \t switch or default (which covers all heading levels)
                if '\\o' not in toc_instr and '\\t' not in toc_instr:
                    # Default TOC without \o includes all heading levels
                    print(f"PASS: Component 2 — TOC uses default levels (all headings) (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 2 — Could not parse TOC level range from: {toc_instr!r}")
        else:
            print(f"FAIL: Component 2 — No TOC field found, cannot check levels")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: A TOC heading/title paragraph exists near the top of the document (0.20 points)
    # The golden doc has a "Table of Contents" Heading 1 near the beginning.
    # We look for any paragraph with "table of contents" or "toc" (case-insensitive)
    # in heading style within the first 10 paragraphs.
    try:
        toc_title_matches = [(i, p) for i, p in enumerate(doc.paragraphs[:10])
                             if ('table of contents' in p.text.strip().lower()
                                 or 'contents' == p.text.strip().lower())
                             and p.style and 'Heading' in p.style.name]
        if len(toc_title_matches) > 0:
            idx, para = toc_title_matches[0]
            total_score += 0.20
            print(f"PASS: Component 3 — TOC title heading at [{idx}]: "
                  f"style={para.style.name!r}, text={para.text!r} (0.20 pts)")
        else:
            print(f"FAIL: Component 3 — No TOC title heading found in first 10 paragraphs")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: TOC is positioned before the main body content (0.20 points)
    # The first original Heading 1 in the initial doc was "Production Engineering".
    # In the golden doc, the TOC content should appear BEFORE "Production Engineering".
    # We check that the TOC field code's paragraph index is before "Production Engineering".
    try:
        toc_para_idx = None
        production_idx = None
        for i, p in enumerate(doc.paragraphs):
            # Find first paragraph containing TOC field code or TOC title
            if toc_para_idx is None:
                xml = p._element.xml
                if 'instrText' in xml and 'TOC' in xml.upper():
                    toc_para_idx = i
                elif p.text.strip().lower() in ('table of contents', 'contents') and \
                     p.style and 'Heading' in p.style.name:
                    toc_para_idx = i
            # Find "Production Engineering" heading
            if production_idx is None and p.style and p.style.name == 'Heading 1' \
               and 'production engineering' in p.text.lower():
                production_idx = i

        if toc_para_idx is not None and production_idx is not None:
            if toc_para_idx < production_idx:
                print(f"PASS: Component 4 — TOC (para {toc_para_idx}) is before "
                      f"main content 'Production Engineering' (para {production_idx}) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — TOC (para {toc_para_idx}) is NOT before "
                      f"'Production Engineering' (para {production_idx})")
        elif toc_para_idx is None:
            print(f"FAIL: Component 4 — Could not locate TOC in document")
        else:
            print(f"FAIL: Component 4 — Could not locate 'Production Engineering' heading")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
