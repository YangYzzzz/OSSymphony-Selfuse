"""
Initial Setup: Create a technical manual document with 6 sections (simulating subdocuments),
each containing Heading 1, Heading 2, and Heading 3 levels. No TOC exists.
Task ID: writer_rm_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_SECTION_START

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Document title (not a heading level that goes into TOC - use Title style)
    title_para = doc.add_heading('Technical Operations Manual', level=0)

    doc.add_paragraph(
        'This master document consolidates all operational procedures, safety protocols, '
        'and technical specifications for the Meridian Advanced Manufacturing Facility. '
        'Each section corresponds to a department-specific subdocument maintained by the '
        'respective department leads.'
    )

    doc.add_paragraph(
        'Document Version: 4.2 | Last Updated: March 2025 | Classification: Internal Use Only'
    )

    # --- Section 1: Production Engineering ---
    doc.add_heading('Production Engineering', level=1)
    doc.add_paragraph(
        'The Production Engineering department oversees all manufacturing processes, '
        'equipment maintenance schedules, and quality assurance protocols for the facility.'
    )

    doc.add_heading('Assembly Line Configuration', level=2)
    doc.add_paragraph(
        'The primary assembly line operates on a three-shift rotation with automated '
        'quality checkpoints at stations 4, 7, and 12. Each station is equipped with '
        'high-resolution cameras and laser measurement systems capable of detecting '
        'deviations as small as 0.002mm from specification tolerances.'
    )

    doc.add_heading('Station Layout and Equipment', level=3)
    doc.add_paragraph(
        'Station 1 through 4 handle raw material preparation, including CNC machining, '
        'surface treatment, and dimensional verification. Stations 5 through 8 perform '
        'subassembly integration with torque-controlled fastening systems.'
    )

    doc.add_heading('Throughput Optimization', level=3)
    doc.add_paragraph(
        'Current throughput targets are set at 847 units per shift with a defect rate '
        'below 0.3%. The bottleneck analysis conducted in Q4 2024 identified station 7 '
        'as the primary constraint, leading to the installation of dual-arm robotic '
        'assistants that increased throughput by 12%.'
    )

    doc.add_heading('Quality Control Procedures', level=2)
    doc.add_paragraph(
        'All finished products undergo a three-stage inspection process: automated optical '
        'inspection (AOI), functional testing under simulated operating conditions, and '
        'final manual inspection by certified quality engineers.'
    )

    doc.add_heading('Defect Classification System', level=3)
    doc.add_paragraph(
        'Defects are categorized into three severity levels: Critical (immediate line stop), '
        'Major (quarantine and review), and Minor (logged for trend analysis). All Critical '
        'defects trigger a root cause analysis within 24 hours.'
    )

    # --- Section 2: Safety and Compliance ---
    doc.add_heading('Safety and Compliance', level=1)
    doc.add_paragraph(
        'The Safety and Compliance department ensures all operations meet OSHA standards, '
        'EPA environmental regulations, and internal corporate safety policies.'
    )

    doc.add_heading('Hazardous Materials Handling', level=2)
    doc.add_paragraph(
        'The facility maintains a comprehensive chemical inventory system tracking 347 '
        'registered substances. Safety Data Sheets (SDS) are accessible at all workstations '
        'through the digital safety portal and physical binder stations.'
    )

    doc.add_heading('Storage Requirements', level=3)
    doc.add_paragraph(
        'Flammable materials are stored in dedicated fire-rated cabinets with automatic '
        'suppression systems. Corrosive substances are segregated in secondary containment '
        'areas with acid-resistant flooring and continuous air monitoring.'
    )

    doc.add_heading('Spill Response Procedures', level=3)
    doc.add_paragraph(
        'In the event of a chemical spill, the first responder must activate the area '
        'alarm, establish a 15-meter exclusion zone, and contact the Emergency Response '
        'Team at extension 5555. Spill kits are located at 23 stations throughout the facility.'
    )

    doc.add_heading('Emergency Response Protocols', level=2)
    doc.add_paragraph(
        'Emergency evacuation routes are reviewed quarterly. Assembly points are designated '
        'at four locations in the parking structure. Full facility drills are conducted '
        'twice annually with participation from the local fire department.'
    )

    doc.add_heading('Fire Suppression Systems', level=3)
    doc.add_paragraph(
        'The facility uses a zoned fire suppression system with FM-200 clean agent in '
        'electronics areas and wet sprinklers in general manufacturing zones. Inspection '
        'and testing follow NFPA 25 requirements on a quarterly basis.'
    )

    # --- Section 3: Information Technology ---
    doc.add_heading('Information Technology', level=1)
    doc.add_paragraph(
        'The IT department manages all digital infrastructure, manufacturing execution '
        'systems (MES), enterprise resource planning (ERP) integration, and cybersecurity.'
    )

    doc.add_heading('Network Infrastructure', level=2)
    doc.add_paragraph(
        'The facility operates on a segmented network architecture with separate VLANs '
        'for corporate operations, manufacturing control systems, and IoT sensor networks. '
        'Redundant fiber connections provide 10 Gbps backbone connectivity.'
    )

    doc.add_heading('Server Architecture', level=3)
    doc.add_paragraph(
        'On-premise servers include a VMware vSphere cluster with 12 ESXi hosts providing '
        '768 GB total RAM and 200 TB SAN storage. Critical workloads are replicated to '
        'the disaster recovery site in real-time via asynchronous replication.'
    )

    doc.add_heading('Wireless Coverage', level=3)
    doc.add_paragraph(
        'Wi-Fi 6E access points are deployed at 45-meter intervals throughout the '
        'production floor, providing minimum 500 Mbps throughput for mobile devices '
        'and handheld scanners used by quality inspectors.'
    )

    doc.add_heading('Cybersecurity Framework', level=2)
    doc.add_paragraph(
        'The facility follows the NIST Cybersecurity Framework with additional controls '
        'specific to industrial control systems (ICS). Network segmentation prevents '
        'lateral movement between IT and OT environments.'
    )

    doc.add_heading('Access Control Policies', level=3)
    doc.add_paragraph(
        'Multi-factor authentication is required for all administrative access. Role-based '
        'access control (RBAC) is enforced across all systems with quarterly access reviews '
        'conducted by department managers and the security team.'
    )

    # --- Section 4: Facilities Management ---
    doc.add_heading('Facilities Management', level=1)
    doc.add_paragraph(
        'Facilities Management is responsible for building maintenance, HVAC systems, '
        'utility management, and space planning for the 42,000 square meter complex.'
    )

    doc.add_heading('HVAC and Environmental Controls', level=2)
    doc.add_paragraph(
        'Clean room areas maintain ISO Class 7 conditions with temperature controlled at '
        '21 plus or minus 1 degree Celsius and relative humidity at 45 plus or minus 5 percent. '
        'The building management system (BMS) monitors 2,300 environmental sensors.'
    )

    doc.add_heading('Air Filtration Systems', level=3)
    doc.add_paragraph(
        'HEPA filters in clean room zones are replaced on a 6-month cycle. Pre-filters '
        'are changed monthly. Differential pressure monitoring ensures filter integrity '
        'with alarms triggered at 1.5 inches of water column pressure drop.'
    )

    doc.add_heading('Temperature Monitoring', level=3)
    doc.add_paragraph(
        'Critical process areas use redundant temperature sensors with data logged at '
        '30-second intervals. Historical data is retained for 7 years to support regulatory '
        'compliance audits and process validation activities.'
    )

    doc.add_heading('Preventive Maintenance Program', level=2)
    doc.add_paragraph(
        'The computerized maintenance management system (CMMS) tracks over 4,500 assets '
        'with scheduled preventive maintenance tasks. Current PM compliance rate is 94.7% '
        'against a target of 95%.'
    )

    doc.add_heading('Equipment Lifecycle Management', level=3)
    doc.add_paragraph(
        'Major equipment undergoes condition-based monitoring using vibration analysis, '
        'thermography, and oil analysis. Replacement decisions are based on total cost of '
        'ownership models updated annually during budget planning.'
    )

    # --- Section 5: Supply Chain and Logistics ---
    doc.add_heading('Supply Chain and Logistics', level=1)
    doc.add_paragraph(
        'The Supply Chain department manages procurement, inventory control, warehouse '
        'operations, and outbound logistics for 1,200 active SKUs across 47 suppliers.'
    )

    doc.add_heading('Inventory Management', level=2)
    doc.add_paragraph(
        'The warehouse operates a hybrid ABC-XYZ classification system. A-class items '
        '(top 20% by value) are cycle-counted weekly, while C-class items are counted '
        'quarterly. The current inventory accuracy rate is 99.2%.'
    )

    doc.add_heading('Warehouse Layout Optimization', level=3)
    doc.add_paragraph(
        'The 8,500 square meter warehouse uses a directed putaway algorithm that assigns '
        'storage locations based on item velocity, size, and pick frequency. Golden zone '
        'shelving (waist to shoulder height) is reserved for the top 50 SKUs by pick volume.'
    )

    doc.add_heading('Automated Storage Systems', level=3)
    doc.add_paragraph(
        'Two vertical lift modules (VLMs) handle small parts storage with a combined '
        'capacity of 12,000 tray positions. Integration with the WMS enables goods-to-person '
        'picking that has reduced average pick time by 40%.'
    )

    doc.add_heading('Supplier Relationship Management', level=2)
    doc.add_paragraph(
        'Suppliers are evaluated quarterly on quality, delivery, cost, and responsiveness '
        'using a weighted scorecard. Strategic suppliers participate in quarterly business '
        'reviews with joint improvement initiatives.'
    )

    doc.add_heading('Vendor Qualification Process', level=3)
    doc.add_paragraph(
        'New vendors undergo a four-stage qualification process: document review, on-site '
        'audit, sample evaluation, and trial production run. Average qualification timeline '
        'is 90 days for non-critical components and 180 days for critical items.'
    )

    # --- Section 6: Human Resources and Training ---
    doc.add_heading('Human Resources and Training', level=1)
    doc.add_paragraph(
        'The HR department manages workforce planning, talent acquisition, employee '
        'development, and compliance training for 623 full-time employees and 85 contractors.'
    )

    doc.add_heading('Training and Certification Programs', level=2)
    doc.add_paragraph(
        'All new employees complete a 40-hour onboarding program covering safety, quality '
        'systems, and role-specific technical training. Annual training hours average 56 '
        'per employee across all departments.'
    )

    doc.add_heading('Technical Certification Tracks', level=3)
    doc.add_paragraph(
        'The facility maintains internal certification programs for CNC operation (3 levels), '
        'robotic systems maintenance (2 levels), and quality inspection (4 levels). '
        'Certifications require annual renewal through practical and written assessments.'
    )

    doc.add_heading('Safety Training Requirements', level=3)
    doc.add_paragraph(
        'OSHA-mandated training includes hazard communication, lockout/tagout, confined '
        'space entry, and powered industrial truck operation. Department-specific training '
        'addresses chemical handling, electrical safety, and ergonomic practices.'
    )

    doc.add_heading('Performance Management', level=2)
    doc.add_paragraph(
        'The performance review cycle runs semi-annually with mid-year check-ins in July '
        'and annual evaluations in January. The system uses a 5-point competency scale '
        'aligned with departmental KPIs and individual development goals.'
    )

    doc.add_heading('Career Development Framework', level=3)
    doc.add_paragraph(
        'Each role has a defined career ladder with clear competency requirements for '
        'advancement. The tuition reimbursement program supports up to $8,000 annually '
        'for approved degree and certification programs.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
