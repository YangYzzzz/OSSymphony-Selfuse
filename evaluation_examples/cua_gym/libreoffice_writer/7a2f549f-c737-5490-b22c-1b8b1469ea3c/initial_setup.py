"""
Initial Setup: Interview transcript document — paragraphs 4 and 8 are plain block quotes (no border/indent yet)
Task ID: writer_para_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_050'
OUTPUT = f'{WORKDIR}/interview_transcript.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1
    h1 = doc.add_heading('Interview with CEO Michael Torres', level=1)

    # Paragraph 2: subtitle/byline
    p2 = doc.add_paragraph('Conducted by Financial Times, February 2025')

    # Paragraph 3: Q
    p3 = doc.add_paragraph('Q: How do you see the company evolving over the next five years?')

    # Paragraph 4: A (block quote — NO border, NO indent in initial state)
    p4 = doc.add_paragraph(
        'A: We are at an inflection point. The investments we made in AI and machine learning '
        'three years ago are now paying dividends. I expect our revenue from AI-driven products '
        'to surpass traditional software licensing within two years.'
    )

    # Paragraph 5: Q
    p5 = doc.add_paragraph('Q: What about concerns regarding AI replacing human workers?')

    # Paragraph 6: narrative
    p6 = doc.add_paragraph('Torres paused before responding, choosing his words carefully.')

    # Paragraph 7: Q
    p7 = doc.add_paragraph('Q: Can you elaborate on your workforce strategy?')

    # Paragraph 8: A (block quote — NO border, NO indent in initial state)
    p8 = doc.add_paragraph(
        'A: We have committed to retraining 100% of affected employees. Our internal data shows '
        'that workers who complete our AI literacy program are 30% more productive. This is not '
        'about replacement \u2014 it is about augmentation and empowerment.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
