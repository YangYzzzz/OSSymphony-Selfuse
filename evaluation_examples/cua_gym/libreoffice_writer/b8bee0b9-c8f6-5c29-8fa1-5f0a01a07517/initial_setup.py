"""
Initial Setup: Set LibreOffice author to 'User1' and open Research_Review.docx with track changes
Task ID: writer_rm_032
Domain: libreoffice_writer
"""

import os
import re
import shlex
import subprocess
import time

# Install dependencies on VM
subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
REG_FILE = f'{WORKDIR}/.config/libreoffice/4/user/registrymodifications.xcu'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def kill_libreoffice():
    """Kill any running LibreOffice processes so config changes take effect."""
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(2)


def set_user_data(first_name, last_name, initials):
    """Set LibreOffice User Data fields in registrymodifications.xcu."""
    with open(REG_FILE, 'r') as f:
        content = f.read()

    # Update givenname
    content = re.sub(
        r'(<item oor:path="/org\.openoffice\.UserProfile/Data"><prop oor:name="givenname" oor:op="fuse"><value>)[^<]*(</value></prop></item>)',
        rf'\g<1>{first_name}\2',
        content
    )

    # Update sn (surname)
    content = re.sub(
        r'(<item oor:path="/org\.openoffice\.UserProfile/Data"><prop oor:name="sn" oor:op="fuse"><value>)[^<]*(</value></prop></item>)',
        rf'\g<1>{last_name}\2',
        content
    )

    # Update initials
    content = re.sub(
        r'(<item oor:path="/org\.openoffice\.UserProfile/Data"><prop oor:name="initials" oor:op="fuse"><value>)[^<]*(</value></prop></item>)',
        rf'\g<1>{initials}\2',
        content
    )

    with open(REG_FILE, 'w') as f:
        f.write(content)

    print(f'User data set: givenname="{first_name}", sn="{last_name}", initials="{initials}"')


def create_document():
    """Create a realistic Research_Review.docx."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('Research Review: Machine Learning Applications in Climate Modeling', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Dr. Emily Watson, Department of Environmental Science')
    run.font.size = Pt(11)
    run.font.italic = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('March 2026')
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This review examines recent advances in the application of machine learning '
        'techniques to climate modeling and prediction. We analyze over 150 peer-reviewed '
        'studies published between 2020 and 2025, focusing on deep learning architectures, '
        'ensemble methods, and physics-informed neural networks. Our findings indicate that '
        'hybrid approaches combining traditional numerical weather prediction with neural '
        'network components achieve the most promising results for medium-range forecasting.'
    )

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Climate modeling represents one of the most computationally intensive scientific '
        'endeavors undertaken by the research community. Traditional general circulation '
        'models (GCMs) solve complex systems of partial differential equations on global '
        'grids, requiring substantial computational resources. The emergence of machine '
        'learning offers potential pathways to accelerate these computations while maintaining '
        'or improving predictive accuracy.'
    )
    doc.add_paragraph(
        'Recent breakthroughs in transformer architectures and graph neural networks have '
        'demonstrated remarkable capabilities in pattern recognition across spatiotemporal '
        'datasets. Google DeepMind\'s GraphCast model, for instance, achieved state-of-the-art '
        'performance in 10-day weather forecasting, outperforming the European Centre for '
        'Medium-Range Weather Forecasts (ECMWF) operational model on multiple metrics.'
    )

    # Methodology
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        'We conducted a systematic literature review following PRISMA guidelines. Database '
        'searches were performed across Web of Science, Scopus, and IEEE Xplore using '
        'the following search terms:'
    )
    doc.add_paragraph('Machine learning AND climate modeling', style='List Bullet')
    doc.add_paragraph('Deep learning AND weather prediction', style='List Bullet')
    doc.add_paragraph('Neural network AND atmospheric science', style='List Bullet')
    doc.add_paragraph('Physics-informed AND climate simulation', style='List Bullet')

    doc.add_paragraph(
        'Studies were included if they: (a) applied at least one ML technique to a '
        'climate-related prediction task, (b) reported quantitative evaluation metrics, '
        'and (c) were published in peer-reviewed venues. A total of 153 studies met our '
        'inclusion criteria after screening 847 initial results.'
    )

    # Results table
    doc.add_heading('3. Key Findings', level=1)
    doc.add_paragraph(
        'Table 1 summarizes the performance comparison across the major architecture '
        'categories identified in our review:'
    )

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Architecture', 'Studies (n)', 'Avg. RMSE Improvement', 'Best Use Case']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['Transformer-based', '42', '23.5%', 'Medium-range (3-10 day)'],
        ['Graph Neural Networks', '31', '18.2%', 'Regional downscaling'],
        ['CNN-based', '38', '15.7%', 'Precipitation nowcasting'],
        ['Physics-Informed NN', '27', '28.1%', 'Long-term projections'],
        ['Ensemble/Hybrid', '15', '31.4%', 'Seasonal forecasting'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    # Discussion
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph(
        'Our analysis reveals several important trends in the field. First, purely '
        'data-driven approaches, while achieving impressive short-term forecasting results, '
        'tend to struggle with climate projections beyond seasonal timescales. This is '
        'consistent with theoretical expectations: without explicit physical constraints, '
        'neural networks may learn correlations that do not generalize under novel forcing '
        'scenarios such as unprecedented greenhouse gas concentrations.'
    )
    doc.add_paragraph(
        'Second, physics-informed neural networks (PINNs) show the highest improvement '
        'rates for long-term projections. By incorporating conservation laws and known '
        'physical relationships as soft constraints in the loss function, these models '
        'maintain physical consistency even when extrapolating beyond the training distribution. '
        'The work of Beucler et al. (2024) demonstrated that PINNs can enforce energy '
        'conservation to within 0.1 W/m² while maintaining competitive forecast skill.'
    )

    # Conclusion
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        'Machine learning is transforming climate science, but the most effective applications '
        'are those that complement rather than replace traditional physical modeling. Hybrid '
        'approaches that leverage the strengths of both paradigms represent the most promising '
        'direction for future research. Key priorities include improving interpretability of '
        'ML models, developing standardized benchmarking protocols, and addressing the '
        'computational cost of training large foundation models for climate applications.'
    )

    # References
    doc.add_heading('References', level=1)
    refs = [
        'Beucler, T., et al. (2024). Enforcing analytic constraints in neural networks emulating physical systems. Physical Review Letters, 132(4), 041501.',
        'Lam, R., et al. (2023). Learning skillful medium-range global weather forecasting. Science, 382(6677), 1416-1421.',
        'Pathak, J., et al. (2022). FourCastNet: A global data-driven high-resolution weather model. arXiv preprint arXiv:2202.11214.',
        'Rasp, S., & Thuerey, N. (2021). Data-driven medium-range weather prediction with a Resnet pretrained on climate simulations. Journal of Advances in Modeling Earth Systems, 13(2).',
        'Weyn, J. A., et al. (2020). Improving data-driven global weather prediction using deep convolutional neural networks on a cubed sphere. Journal of Advances in Modeling Earth Systems, 12(9).',
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f'[{i}] {ref}')

    doc.save(OUTPUT)
    print(f'Document created: {OUTPUT}')


def main():
    # Kill LibreOffice first so config changes take effect
    kill_libreoffice()

    # Set author to 'User1' in LibreOffice user data
    set_user_data('User1', '', 'U1')

    # Create the research document
    create_document()

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)

    # Enable track changes via macro after app starts
    # We use a brief delay then send the macro command
    time.sleep(2)
    # Enable track changes using command line macro dispatch
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        ['python3', '-c', '''
import subprocess, time
time.sleep(3)
# Use xdotool to enable track changes via keyboard shortcut Ctrl+Shift+C
import os
env = os.environ.copy()
env["DISPLAY"] = ":0"
subprocess.run(["xdotool", "key", "ctrl+shift+c"], env=env)
'''],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )

    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
    print('Track changes will be enabled via Ctrl+Shift+C shortcut')


main()
