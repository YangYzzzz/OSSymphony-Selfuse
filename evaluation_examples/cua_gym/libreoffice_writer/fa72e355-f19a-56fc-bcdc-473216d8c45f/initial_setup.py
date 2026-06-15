"""
Initial Setup: Climate literature review document with 5 bibliography references.
The conclusion section contains a placeholder '(cite here)' that needs to be replaced.
Task ID: osworld_writer_biblio_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_biblio_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ---- Title ----
    title_para = doc.add_heading('Climate Change Mitigation and Adaptation: A Literature Review', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    author_para = doc.add_paragraph('Prepared by: Environmental Research Group')
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para = doc.add_paragraph('March 2025')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # blank spacer

    # ---- Abstract ----
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This literature review synthesizes current research on climate change mitigation and adaptation '
        'strategies. Drawing on peer-reviewed publications from 2018 to 2024, we examine the effectiveness '
        'of policy interventions, technological innovations, and community-based approaches. The review '
        'identifies significant gaps in the literature regarding long-term monitoring frameworks and '
        'cross-sector collaboration mechanisms.'
    )

    # ---- Introduction ----
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Climate change represents one of the most pressing challenges of the twenty-first century. '
        'Rising global temperatures, increasingly frequent extreme weather events, and accelerating '
        'sea-level rise threaten both natural ecosystems and human societies worldwide. The scientific '
        'consensus on anthropogenic climate change is unequivocal, with the Intergovernmental Panel on '
        'Climate Change (IPCC) reporting that global mean surface temperature has already risen by '
        'approximately 1.1°C above pre-industrial levels (Smith et al., 2021).'
    )
    doc.add_paragraph(
        'In response to these challenges, researchers and policymakers have developed a broad spectrum '
        'of mitigation and adaptation strategies. Mitigation efforts focus on reducing greenhouse gas '
        'emissions through renewable energy transitions, energy efficiency improvements, and land-use '
        'changes. Adaptation strategies, by contrast, aim to adjust societal and ecological systems to '
        'actual or anticipated climate change and its effects (Brown & Martinez, 2020).'
    )
    doc.add_paragraph(
        'This review examines the peer-reviewed literature published between 2018 and 2024 with a '
        'focus on three thematic areas: (1) renewable energy deployment and grid integration, '
        '(2) carbon capture and sequestration technologies, and (3) urban climate adaptation frameworks. '
        'The objectives of this review are to identify key findings, methodological approaches, and '
        'remaining research gaps in each thematic area.'
    )

    # ---- Section 2 ----
    doc.add_heading('2. Renewable Energy Deployment', level=1)
    doc.add_paragraph(
        'The transition to renewable energy sources has accelerated dramatically over the past decade. '
        'Solar photovoltaic and wind energy capacity have grown at compound annual rates exceeding 20%, '
        'driven by dramatic cost reductions and supportive policy environments in major economies '
        '(Thompson & Lee, 2023). In many regions, newly installed renewable capacity is now less '
        'expensive than continued operation of existing fossil fuel plants.'
    )
    doc.add_paragraph(
        'Grid integration challenges have emerged as a critical bottleneck for high-penetration renewable '
        'systems. Variable generation from wind and solar requires complementary flexibility resources, '
        'including grid-scale battery storage, demand response programs, and enhanced transmission '
        'infrastructure. Recent modeling studies suggest that electricity systems with 80–90% renewable '
        'penetration are technically and economically feasible with appropriate investment in storage '
        'and grid management technologies (Chen et al., 2022).'
    )
    doc.add_paragraph(
        'Emerging research has also highlighted significant equity dimensions of the energy transition. '
        'Communities historically dependent on fossil fuel industries face substantial economic '
        'disruption, raising important questions about just transition pathways. Thompson and Lee (2023) '
        'found that proactive workforce retraining programs and economic diversification initiatives '
        'significantly reduced adverse employment impacts in transitioning regions.'
    )

    # ---- Section 3 ----
    doc.add_heading('3. Carbon Capture and Sequestration', level=1)
    doc.add_paragraph(
        'Carbon capture and sequestration (CCS) technologies have attracted growing attention as a '
        'potential complement to direct emissions reductions. Industrial CCS facilities capture CO₂ '
        'at the point of emission from power plants and heavy industry, compressing and injecting it '
        'into geological formations for long-term storage. As of 2024, approximately 45 large-scale '
        'CCS projects are operational worldwide, with a combined capture capacity of around 50 million '
        'tonnes of CO₂ per year (Chen et al., 2022).'
    )
    doc.add_paragraph(
        'Bioenergy with carbon capture and storage (BECCS) represents a potentially negative-emission '
        'approach, combining biomass combustion with geological CO₂ storage. However, land-use '
        'requirements for bioenergy feedstocks raise significant concerns about competition with food '
        'production and biodiversity conservation. Multiple studies have cautioned that large-scale '
        'BECCS deployment could have adverse environmental and social consequences if implemented '
        'without robust governance frameworks.'
    )
    doc.add_paragraph(
        'Direct air capture (DAC) technologies, which extract CO₂ directly from the atmosphere, '
        'offer greater geographic flexibility than point-source CCS but currently require substantially '
        'more energy input. Ongoing research and development efforts are focused on reducing energy '
        'requirements and material costs, with several pilot facilities demonstrating incremental '
        'technical progress over the past five years.'
    )

    # ---- Section 4 ----
    doc.add_heading('4. Urban Climate Adaptation', level=1)
    doc.add_paragraph(
        'Cities face disproportionate climate risks due to the concentration of people, infrastructure, '
        'and economic activity in areas vulnerable to flooding, heat stress, and other climate hazards. '
        'Urban heat island effects amplify thermal stress beyond regional temperature increases, '
        'disproportionately affecting low-income communities with limited access to cooling resources. '
        'Kim and Patel (2019) documented mean summer temperature differences of 3–7°C between urban '
        'core areas and surrounding rural regions across a global sample of 50 major metropolitan areas.'
    )
    doc.add_paragraph(
        'Nature-based solutions have gained prominence as cost-effective urban adaptation strategies. '
        'Green infrastructure—including urban forests, green roofs, permeable pavements, and restored '
        'urban waterways—provides multiple co-benefits beyond climate resilience, including improved '
        'air quality, biodiversity support, and psychological wellbeing. A meta-analysis of 73 urban '
        'greening interventions found average surface temperature reductions of 2.1°C within treated '
        'areas, with additional heat-mitigation benefits extending 100–300 metres beyond treatment '
        'boundaries (Kim & Patel, 2019).'
    )
    doc.add_paragraph(
        'Governance frameworks for urban climate adaptation vary widely. Cities in high-income countries '
        'generally possess greater institutional capacity to develop and implement long-term adaptation '
        'plans, while many rapidly urbanizing cities in lower-income contexts face significant resource '
        'and capacity constraints. Brown and Martinez (2020) emphasised the importance of multi-level '
        'governance arrangements that integrate local knowledge with national policy frameworks and '
        'international finance mechanisms.'
    )

    # ---- Conclusion ----
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        'This literature review has examined three interrelated dimensions of climate change response: '
        'renewable energy deployment, carbon capture and sequestration, and urban adaptation. Across '
        'all three areas, the evidence base has grown substantially over the review period, reflecting '
        'both heightened research attention and accelerating real-world implementation experience.'
    )
    doc.add_paragraph(
        'Key findings include the rapid cost decline and technical maturation of renewable energy '
        'technologies, the persistent challenges associated with large-scale CCS deployment, and the '
        'growing evidence base for nature-based solutions in urban environments. Despite this progress, '
        'critical research gaps remain, particularly regarding long-term monitoring of CCS storage '
        'integrity, equity dimensions of energy transitions, and climate adaptation in rapidly '
        'urbanizing low-income contexts.'
    )
    concl_para = doc.add_paragraph(
        'Recent research has begun to explore integrated approaches that combine mitigation and '
        'adaptation strategies at the city scale, recognising that these are not separate challenges '
        'but deeply interconnected dimensions of a coherent climate response '
        '(cite here)'
        '. Future work should prioritise multi-sector, multi-scale analyses that capture these '
        'interactions and their implications for policy design.'
    )
    doc.add_paragraph(
        'The urgency of the climate challenge demands continued expansion and deepening of the '
        'research base. Interdisciplinary collaboration, open data sharing, and close engagement '
        'between researchers and policymakers will be essential to translate scientific knowledge '
        'into effective climate action.'
    )

    # ---- References ----
    doc.add_heading('References', level=1)

    refs = [
        ('1. ', 'Brown, K. L., & Martinez, J. R. (2020). Governing urban climate adaptation: '
                'Multi-level frameworks and local implementation. '
                'Global Environmental Change, 62, 102068. '
                'https://doi.org/10.1016/j.gloenvcha.2020.102068'),
        ('2. ', 'Chen, W., Liu, Y., & Nakamura, T. (2022). High-penetration renewable electricity '
                'systems: Grid flexibility, storage, and cost trajectories to 2050. '
                'Nature Energy, 7(3), 234–246. '
                'https://doi.org/10.1038/s41560-022-00987-3'),
        ('3. ', 'Kim, S. J., & Patel, R. V. (2019). Urban heat islands and green infrastructure: '
                'A global meta-analysis of mitigation potential. '
                'Landscape and Urban Planning, 185, 103736. '
                'https://doi.org/10.1016/j.landurbplan.2019.103736'),
        ('4. ', 'Smith, A. D., Foster, E. M., & Okonkwo, C. (2021). Global surface temperature '
                'trends: Updated analysis incorporating ocean heat content adjustments. '
                'Journal of Climate, 34(11), 4305–4322. '
                'https://doi.org/10.1175/JCLI-D-20-0512.1'),
        ('5. ', 'Thompson, R. G., & Lee, H. S. (2023). Just transition pathways in fossil '
                'fuel-dependent regions: Evidence from 12 case studies. '
                'Energy Research & Social Science, 96, 102913. '
                'https://doi.org/10.1016/j.erss.2022.102913'),
    ]

    for num, text in refs:
        para = doc.add_paragraph()
        run_num = para.add_run(num)
        run_num.bold = False
        run_text = para.add_run(text)
        run_text.bold = False
        # Hanging indent style for references
        para.paragraph_format.left_indent = Pt(36)
        para.paragraph_format.first_line_indent = Pt(-36)
        para.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
