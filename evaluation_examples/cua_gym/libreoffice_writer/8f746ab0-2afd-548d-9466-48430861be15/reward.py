"""
Reward Script: Add confidential header to HR policy document
Task ID: writer_hr_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Header is active (not empty, not linked_to_previous) in section 0
  Component 2 (0.4): Header text matches 'CONFIDENTIAL - INTERNAL USE ONLY'
  Component 3 (0.2): Header is present in ALL sections (covers every page)
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_004'
EXPECTED_HEADER = 'CONFIDENTIAL - INTERNAL USE ONLY'


def persist_app_state(domain: str):
    """Best-effort save of any unsaved GUI edits."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = doc.sections
    if len(sections) == 0:
        print("FAIL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Header is active in section 0 (0.4 points)
    # The header must not be empty and should be explicitly set (not linked_to_previous
    # with empty content). This checks that a header was actually added.
    try:
        header = sections[0].header
        header_text = ''
        for para in header.paragraphs:
            header_text += para.text
        header_text = header_text.strip()

        if header_text:
            print(f"PASS: Component 1 — Header is active with text: '{header_text}' (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Header is empty or not set")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header text matches expected value exactly (0.4 points)
    # The task requires the exact text 'CONFIDENTIAL - INTERNAL USE ONLY'
    try:
        header = sections[0].header
        header_text = ''
        for para in header.paragraphs:
            header_text += para.text
        header_text = header_text.strip()

        if header_text == EXPECTED_HEADER:
            print(f"PASS: Component 2 — Header text matches exactly (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Expected '{EXPECTED_HEADER}', found '{header_text}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header is present in ALL sections (0.2 points)
    # For a multi-section document, every section must show the header.
    # For a single-section document, this is satisfied if section 0 has it.
    try:
        all_sections_have_header = True
        for i, section in enumerate(sections):
            sec_header = section.header
            sec_text = ''
            for para in sec_header.paragraphs:
                sec_text += para.text
            sec_text = sec_text.strip()

            if EXPECTED_HEADER not in sec_text:
                all_sections_have_header = False
                print(f"FAIL: Component 3 — Section {i} header missing expected text, found: '{sec_text}'")
                break

        if all_sections_have_header:
            print(f"PASS: Component 3 — All {len(sections)} section(s) have the header (0.2 pts)")
            total_score += 0.2
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved GUI state before verification
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
