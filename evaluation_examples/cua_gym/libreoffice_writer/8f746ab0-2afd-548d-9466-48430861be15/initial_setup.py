"""
Initial Setup: Add a header to every page of an HR policy document
Task ID: writer_hr_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Ensure NO header exists
    section.header.is_linked_to_previous = True
    section.different_first_page_header_footer = False

    # --- Title Page ---
    title = doc.add_heading('HR Policies & Procedures Manual 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Global Solutions Inc.')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info.add_run('Effective Date: January 1, 2026\nRevision: 4.2\nApproved by: Elena Vasquez, Chief Human Resources Officer')
    run.font.size = Pt(11)

    doc.add_page_break()

    # --- Page 2: Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction and Purpose ............................ 3',
        '2. Equal Employment Opportunity ....................... 3',
        '3. Code of Conduct ................................... 4',
        '4. Compensation and Benefits ......................... 5',
        '5. Leave Policies .................................... 6',
        '6. Workplace Safety and Health ........................ 7',
        '7. Disciplinary Procedures ........................... 7',
        '8. Acknowledgment and Acceptance ..................... 8',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # --- Page 3: Section 1 - Introduction ---
    doc.add_heading('1. Introduction and Purpose', level=1)
    doc.add_paragraph(
        'This Human Resources Policies and Procedures Manual has been developed to provide '
        'all employees of Meridian Global Solutions Inc. with a comprehensive guide to company '
        'policies, workplace expectations, and available benefits. It is designed to promote a '
        'fair, consistent, and productive work environment for every team member.'
    )
    doc.add_paragraph(
        'All employees are expected to familiarize themselves with the contents of this manual '
        'and to comply with the policies and procedures outlined herein. This manual supersedes '
        'all previous policy documents and memoranda. Questions regarding interpretation or '
        'application of any policy should be directed to the Human Resources Department at '
        'hr@meridianglobal.com or extension 4500.'
    )
    doc.add_paragraph(
        'Meridian Global Solutions reserves the right to modify, revoke, suspend, or terminate '
        'any policy at any time, with or without notice. Updated versions will be distributed '
        'electronically and posted on the company intranet portal.'
    )

    # --- Section 2: Equal Employment Opportunity ---
    doc.add_heading('2. Equal Employment Opportunity', level=1)
    doc.add_paragraph(
        'Meridian Global Solutions is committed to providing equal employment opportunities '
        'to all qualified individuals regardless of race, color, religion, sex, sexual orientation, '
        'gender identity, national origin, age, disability, genetic information, veteran status, '
        'or any other protected characteristic under applicable federal, state, or local law.'
    )
    doc.add_paragraph(
        'This commitment extends to all aspects of employment, including recruitment, hiring, '
        'placement, promotion, transfer, training, compensation, benefits, social and recreational '
        'programs, and termination. Any employee who believes they have been subjected to '
        'discrimination should report the matter immediately to their supervisor, the HR '
        'Department, or through the anonymous ethics hotline at 1-800-555-0147.'
    )
    doc.add_paragraph(
        'Retaliation against any individual who files a complaint or participates in an '
        'investigation is strictly prohibited and will result in disciplinary action up to '
        'and including termination of employment.'
    )

    doc.add_page_break()

    # --- Page 4: Section 3 - Code of Conduct ---
    doc.add_heading('3. Code of Conduct', level=1)
    doc.add_paragraph(
        'All employees are expected to conduct themselves with integrity, professionalism, '
        'and respect for colleagues, clients, and business partners. The following standards '
        'of conduct apply to all employees at every level of the organization:'
    )

    conduct_items = [
        'Maintain honesty and transparency in all business dealings and communications.',
        'Treat all individuals with dignity, courtesy, and respect regardless of position or background.',
        'Protect confidential company information, trade secrets, and proprietary data from unauthorized disclosure.',
        'Avoid conflicts of interest and disclose any potential conflicts to your supervisor promptly.',
        'Comply with all applicable laws, regulations, and company policies at all times.',
        'Use company resources, including technology and equipment, responsibly and for authorized purposes only.',
        'Report any suspected violations of company policy, unethical behavior, or illegal activity.',
        'Maintain a professional appearance and demeanor consistent with your role and department standards.',
    ]
    for item in conduct_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Violations of the Code of Conduct may result in disciplinary action, including verbal '
        'or written warnings, suspension, demotion, or termination, depending on the severity '
        'and nature of the violation. The company reserves the right to involve law enforcement '
        'authorities when appropriate.'
    )

    doc.add_page_break()

    # --- Page 5: Section 4 - Compensation and Benefits ---
    doc.add_heading('4. Compensation and Benefits', level=1)

    doc.add_heading('4.1 Salary Structure', level=2)
    doc.add_paragraph(
        'Meridian Global Solutions maintains a competitive compensation structure based on '
        'market analysis, internal equity, and individual performance. Salary ranges are '
        'reviewed annually by the Compensation Committee and adjusted as necessary to remain '
        'competitive within our industry sector.'
    )

    doc.add_heading('4.2 Performance Reviews', level=2)
    doc.add_paragraph(
        'Formal performance evaluations are conducted semi-annually in June and December. '
        'Employees receive written feedback on their performance, goal achievement, and areas '
        'for professional development. Merit-based salary increases are determined during the '
        'annual review cycle and take effect on April 1 of each fiscal year.'
    )

    doc.add_heading('4.3 Benefits Overview', level=2)
    doc.add_paragraph(
        'Full-time employees working 30 or more hours per week are eligible for the following benefits:'
    )
    benefits = [
        'Medical insurance (PPO and HMO options) with 80% employer contribution',
        'Dental and vision insurance with 70% employer contribution',
        '401(k) retirement plan with 6% employer match (vesting over 3 years)',
        'Life insurance at 2x annual salary, provided at no cost to the employee',
        'Short-term and long-term disability coverage',
        'Employee Assistance Program (EAP) providing counseling and support services',
        'Tuition reimbursement up to $5,250 per calendar year for approved programs',
        'Professional development budget of $1,500 per year for conferences and certifications',
    ]
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # --- Page 6: Section 5 - Leave Policies ---
    doc.add_heading('5. Leave Policies', level=1)

    doc.add_heading('5.1 Paid Time Off (PTO)', level=2)
    doc.add_paragraph(
        'PTO accrual is based on length of service as follows:'
    )

    # PTO table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['Years of Service', 'Annual PTO Days', 'Monthly Accrual']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    pto_data = [
        ['0 - 2 years', '15 days', '1.25 days'],
        ['3 - 7 years', '20 days', '1.67 days'],
        ['8+ years', '25 days', '2.08 days'],
    ]
    for r, row_data in enumerate(pto_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    doc.add_heading('5.2 Sick Leave', level=2)
    doc.add_paragraph(
        'Employees accrue 10 sick days per year, accrued monthly. Unused sick leave carries '
        'over to the following year up to a maximum of 60 days. A medical certificate is '
        'required for absences exceeding three consecutive workdays.'
    )

    doc.add_heading('5.3 Parental Leave', level=2)
    doc.add_paragraph(
        'Meridian Global Solutions provides 16 weeks of paid parental leave for the birth '
        'or adoption of a child. This applies to all eligible employees regardless of gender. '
        'Additional unpaid leave of up to 12 weeks may be requested under the Family and '
        'Medical Leave Act (FMLA).'
    )

    doc.add_page_break()

    # --- Page 7: Section 6 - Workplace Safety ---
    doc.add_heading('6. Workplace Safety and Health', level=1)
    doc.add_paragraph(
        'Meridian Global Solutions is committed to maintaining a safe and healthy work '
        'environment for all employees, contractors, and visitors. The company complies with '
        'all applicable Occupational Safety and Health Administration (OSHA) regulations and '
        'industry best practices.'
    )
    doc.add_paragraph(
        'All employees are responsible for following safety procedures, reporting hazards, '
        'and participating in required safety training programs. Workplace injuries, no matter '
        'how minor, must be reported to the employee\'s supervisor and the Safety Committee '
        'within 24 hours of occurrence. Incident report forms are available on the company '
        'intranet and from the HR Department.'
    )
    doc.add_paragraph(
        'The company maintains an Emergency Action Plan that includes evacuation procedures, '
        'severe weather protocols, and active threat response guidelines. Fire drills are '
        'conducted quarterly and participation is mandatory for all on-site personnel.'
    )

    # --- Section 7: Disciplinary Procedures ---
    doc.add_heading('7. Disciplinary Procedures', level=1)
    doc.add_paragraph(
        'Meridian Global Solutions follows a progressive discipline policy designed to '
        'provide employees with the opportunity to correct performance or behavioral issues. '
        'The progressive discipline process consists of the following steps:'
    )

    steps = [
        'Step 1: Verbal Warning - A documented conversation between the supervisor and employee addressing the issue and expected improvement.',
        'Step 2: Written Warning - A formal written notice detailing the problem, previous discussions, and specific corrective actions required within a defined timeframe.',
        'Step 3: Final Written Warning / Suspension - A final formal notice that may include unpaid suspension of 1 to 5 days. Failure to improve will result in termination.',
        'Step 4: Termination - Separation from employment when previous corrective actions have not resulted in satisfactory improvement.',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph(
        'The company reserves the right to skip steps in the progressive discipline process '
        'for serious violations, including but not limited to theft, fraud, workplace violence, '
        'harassment, substance abuse on company premises, or gross insubordination.'
    )

    doc.add_page_break()

    # --- Page 8: Section 8 - Acknowledgment ---
    doc.add_heading('8. Acknowledgment and Acceptance', level=1)
    doc.add_paragraph(
        'By signing below, I acknowledge that I have received, read, and understood the '
        'Meridian Global Solutions HR Policies & Procedures Manual 2026. I agree to comply '
        'with all policies and procedures described herein and understand that failure to '
        'do so may result in disciplinary action up to and including termination of employment.'
    )
    doc.add_paragraph(
        'I understand that this manual is not a contract of employment and does not create '
        'any contractual obligations on the part of Meridian Global Solutions Inc. Employment '
        'with the company remains at-will, meaning that either the employee or the company '
        'may terminate the employment relationship at any time, with or without cause or notice.'
    )
    doc.add_paragraph()

    # Signature lines
    sig_fields = [
        'Employee Name (Print): _________________________________________',
        '',
        'Employee Signature: ____________________________________________',
        '',
        'Date: ___________________',
        '',
        'Department: _________________________________________________',
        '',
        'Supervisor Name: _____________________________________________',
        '',
        'Supervisor Signature: __________________________________________',
        '',
        'Date: ___________________',
    ]
    for field in sig_fields:
        p = doc.add_paragraph(field)
        p.paragraph_format.space_after = Pt(2)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
