"""
Reward Script: Change numbered list from '1. 2. 3.' to 'a) b) c)' lettered numbering
Task ID: wrpara_010
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): numFmt is lowerLetter (not decimal)
  Component 2 (0.3): lvlText uses ')' suffix pattern (e.g. '%1)')
  Component 3 (0.2): All 5 list items preserved with numbering references intact
"""

import os
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'wrpara_010'
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WNS}


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_numbering_format(doc_path):
    """
    Extract the numFmt and lvlText for the numbering definition used by
    the list paragraphs (numId=10 -> abstractNumId, level 0).
    Returns (numFmt_val, lvlText_val) or (None, None) on error.
    """
    from docx import Document

    doc = Document(doc_path)
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    # Step 1: Find which abstractNumId is referenced by numId used in the list paragraphs
    # Get the numId from the first list paragraph
    list_num_id = None
    for para in doc.paragraphs:
        numPr = para._element.find('.//w:pPr/w:numPr', NS)
        if numPr is not None:
            numId_el = numPr.find('w:numId', NS)
            if numId_el is not None:
                list_num_id = numId_el.get(f'{{{WNS}}}val')
                break

    if list_num_id is None:
        return None, None

    # Step 2: Find the abstractNumId for this numId
    abstract_num_id = None
    for num_el in numbering_xml.findall('.//w:num', NS):
        if num_el.get(f'{{{WNS}}}numId') == list_num_id:
            abs_ref = num_el.find('w:abstractNumId', NS)
            if abs_ref is not None:
                abstract_num_id = abs_ref.get(f'{{{WNS}}}val')
            # Also check for lvlOverride which can override the format
            lvl_override = num_el.find('.//w:lvlOverride/w:lvl', NS)
            if lvl_override is not None:
                fmt_el = lvl_override.find('w:numFmt', NS)
                txt_el = lvl_override.find('w:lvlText', NS)
                if fmt_el is not None and txt_el is not None:
                    return (fmt_el.get(f'{{{WNS}}}val'),
                            txt_el.get(f'{{{WNS}}}val'))
            break

    if abstract_num_id is None:
        return None, None

    # Step 3: Find the abstractNum and get lvl 0 format
    for abs_el in numbering_xml.findall('.//w:abstractNum', NS):
        if abs_el.get(f'{{{WNS}}}abstractNumId') == abstract_num_id:
            lvl0 = None
            for lvl in abs_el.findall('w:lvl', NS):
                if lvl.get(f'{{{WNS}}}ilvl') == '0':
                    lvl0 = lvl
                    break
            if lvl0 is None:
                return None, None
            fmt_el = lvl0.find('w:numFmt', NS)
            txt_el = lvl0.find('w:lvlText', NS)
            num_fmt = fmt_el.get(f'{{{WNS}}}val') if fmt_el is not None else None
            lvl_text = txt_el.get(f'{{{WNS}}}val') if txt_el is not None else None
            return num_fmt, lvl_text

    return None, None


def count_numbered_paragraphs(doc_path):
    """Count paragraphs that have a numPr (numbering properties)."""
    from docx import Document

    doc = Document(doc_path)
    count = 0
    for para in doc.paragraphs:
        numPr = para._element.find('.//w:pPr/w:numPr', NS)
        if numPr is not None:
            numId_el = numPr.find('w:numId', NS)
            if numId_el is not None:
                val = numId_el.get(f'{{{WNS}}}val')
                # numId=0 means no numbering
                if val and val != '0':
                    count += 1
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must be loadable
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get the numbering format
    try:
        num_fmt, lvl_text = get_numbering_format(file_path)
        print(f"INFO: numFmt={num_fmt!r}, lvlText={lvl_text!r}")
    except Exception as e:
        print(f"ERROR: Could not extract numbering format: {e}")
        num_fmt, lvl_text = None, None

    # Component 1: numFmt is 'lowerLetter' (0.5 points)
    # This is the core change: decimal -> lowerLetter
    try:
        if num_fmt == 'lowerLetter':
            print(f"PASS: Component 1 -- numFmt is 'lowerLetter' (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 -- Expected numFmt='lowerLetter', found: {num_fmt!r}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: lvlText uses ')' suffix pattern like '%1)' (0.3 points)
    # This verifies the parenthesis format: a) b) c) instead of a. b. c.
    try:
        if lvl_text is not None and ')' in lvl_text:
            print(f"PASS: Component 2 -- lvlText contains ')' pattern: {lvl_text!r} (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 -- Expected lvlText with ')' suffix, found: {lvl_text!r}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: All 5 list items preserved AND using lowerLetter format (0.2 points)
    # This is a compound check anchored to the task change: items must be intact
    # AND the numbering format must already be lowerLetter (ties to the actual change)
    try:
        num_count = count_numbered_paragraphs(file_path)
        list_texts = []
        for para in doc.paragraphs:
            numPr = para._element.find('.//w:pPr/w:numPr', NS)
            if numPr is not None:
                numId_el = numPr.find('w:numId', NS)
                if numId_el is not None and numId_el.get(f'{{{WNS}}}val') not in (None, '0'):
                    list_texts.append(para.text.strip())

        # All 5 items preserved with content AND numbering is lowerLetter
        if num_count == 5 and all(len(t) > 10 for t in list_texts) and num_fmt == 'lowerLetter':
            print(f"PASS: Component 3 -- All 5 list items preserved with lowerLetter numbering (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 -- Expected 5 items with lowerLetter. Found {num_count} items, numFmt={num_fmt!r}")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = round(min(total_score, 1.0), 1)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state in case edits are unsaved in LibreOffice
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
