"""
Initial Setup: Create a document with a numbered list using Arabic numeral format (1. 2. 3.)
Task ID: wrpara_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'wrpara_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_numbering_part(doc, num_fmt='decimal', lvl_text='%1.', start=1):
    """Create a numbering definition with the specified format."""
    # Access or create numbering part
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Find the highest abstractNumId already in use
    abstract_nums = numbering_elm.findall(qn('w:abstractNum'))
    max_abstract_id = -1
    for an in abstract_nums:
        aid = int(an.get(qn('w:abstractNumId')))
        if aid > max_abstract_id:
            max_abstract_id = aid
    new_abstract_id = max_abstract_id + 1

    # Create abstractNum element
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(new_abstract_id))

    # Multi-level type
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'hybridMultilevel')
    abstract_num.append(multi)

    # Level 0 definition
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')

    start_elm = OxmlElement('w:start')
    start_elm.set(qn('w:val'), str(start))
    lvl.append(start_elm)

    numFmt = OxmlElement('w:numFmt')
    numFmt.set(qn('w:val'), num_fmt)
    lvl.append(numFmt)

    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), lvl_text)
    lvl.append(lvlText)

    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(qn('w:val'), 'left')
    lvl.append(lvlJc)

    # Indentation
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    lvl.append(pPr)

    abstract_num.append(lvl)

    # Insert abstractNum before the first w:num element
    nums = numbering_elm.findall(qn('w:num'))
    if nums:
        numbering_elm.insert(list(numbering_elm).index(nums[0]), abstract_num)
    else:
        numbering_elm.append(abstract_num)

    # Create num element referencing the abstractNum
    num_ids = numbering_elm.findall(qn('w:num'))
    max_num_id = 0
    for n in num_ids:
        nid = int(n.get(qn('w:numId')))
        if nid > max_num_id:
            max_num_id = nid
    new_num_id = max_num_id + 1

    num_elm = OxmlElement('w:num')
    num_elm.set(qn('w:numId'), str(new_num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(new_abstract_id))
    num_elm.append(abstract_ref)
    numbering_elm.append(num_elm)

    return new_num_id


def add_numbered_paragraph(doc, text, num_id, ilvl=0):
    """Add a paragraph with specific numbering."""
    para = doc.add_paragraph()
    # Set numbering properties
    pPr = para._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_elm = OxmlElement('w:ilvl')
    ilvl_elm.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_elm)
    numId_elm = OxmlElement('w:numId')
    numId_elm.set(qn('w:val'), str(num_id))
    numPr.append(numId_elm)
    pPr.append(numPr)
    # Add text
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return para


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Procedure Steps', level=1)

    # Add a brief intro paragraph
    intro = doc.add_paragraph(
        'The following steps outline the standard procedure for onboarding '
        'new employees at Meridian Technologies. Please follow each step in '
        'order to ensure a smooth transition.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Create numbered list with Arabic numeral format: "1." "2." etc.
    # We need to use the "List Number" style which gives standard 1. 2. 3. numbering
    # But to be explicit about the format, let's create our own numbering definition
    num_id = create_numbering_part(doc, num_fmt='decimal', lvl_text='%1.', start=1)

    # Five numbered items with realistic content
    items = [
        'Submit the completed HR intake form along with two forms of identification '
        'to the Human Resources department by the end of the first business day.',

        'Schedule a meeting with your assigned team lead, Rachel Morrison, to review '
        'the project roadmap and current sprint objectives for Q2 2025.',

        'Complete the mandatory cybersecurity awareness training module available '
        'on the internal learning portal before accessing any production systems.',

        'Set up your development environment by following the configuration guide '
        'located in the shared Engineering wiki under "New Developer Setup".',

        'Attend the weekly all-hands meeting held every Wednesday at 10:00 AM in '
        'Conference Room B to familiarize yourself with ongoing initiatives.',
    ]

    for item_text in items:
        add_numbered_paragraph(doc, item_text, num_id)

    # Add a closing paragraph
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    run = closing.add_run(
        'If you have any questions about these procedures, please contact '
        'the HR Support team at hr-support@meridiantech.com.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
