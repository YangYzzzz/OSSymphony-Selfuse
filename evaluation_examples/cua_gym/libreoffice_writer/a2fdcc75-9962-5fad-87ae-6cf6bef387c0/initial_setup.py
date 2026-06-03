"""
Initial Setup: Fintech Africa Essay with References Section
Task ID: osworld_writer_biblio_009
Domain: libreoffice_writer

Creates a research essay about fintech adoption in Africa with:
- Introduction, 4 body paragraphs, conclusion, and References section
- Third paragraph contains '(ref needed)' placeholder
- References section has 7 APA-format numbered entries (1-7)
- Reference #8 and cross-reference '(8)' are NOT in the initial state
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_biblio_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Fintech Adoption and Financial Inclusion in Sub-Saharan Africa: Challenges and Opportunities")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # blank line

    # --- Author info ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.add_run("A Review of Current Literature and Emerging Trends")

    doc.add_paragraph()  # blank line

    # --- Introduction ---
    intro_heading = doc.add_paragraph()
    run = intro_heading.add_run("Introduction")
    run.bold = True
    run.font.size = Pt(12)

    intro_text = doc.add_paragraph(
        "The rapid proliferation of financial technology (fintech) across sub-Saharan Africa has fundamentally "
        "transformed how millions of people access financial services. Mobile money platforms, digital lending "
        "solutions, and blockchain-based remittance services have emerged as critical enablers of financial "
        "inclusion in a region where traditional banking infrastructure has historically been inadequate (1). "
        "This essay examines the multifaceted relationship between fintech adoption and financial inclusion "
        "across sub-Saharan Africa, drawing on recent empirical research to assess both progress and persistent "
        "challenges."
    )
    intro_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- Body Paragraph 1 ---
    body1_heading = doc.add_paragraph()
    run = body1_heading.add_run("The Landscape of Financial Exclusion")
    run.bold = True
    run.font.size = Pt(12)

    body1_text = doc.add_paragraph(
        "Sub-Saharan Africa contains some of the world's highest rates of financial exclusion, with approximately "
        "57% of adults remaining unbanked as of 2021. Limited physical banking infrastructure, high transaction "
        "costs, stringent documentation requirements, and geographic barriers have historically prevented low-income "
        "populations from accessing formal financial services (2). Rural communities in particular face acute "
        "challenges, as bank branches and ATMs remain concentrated in urban centers, forcing many households to "
        "rely on informal savings mechanisms such as rotating savings and credit associations (ROSCAs) and "
        "moneylenders. The consequences of financial exclusion extend beyond mere inconvenience: excluded "
        "populations face difficulty accumulating savings, accessing credit for productive investment, and "
        "managing income volatility, perpetuating cycles of poverty (3)."
    )
    body1_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- Body Paragraph 2 ---
    body2_heading = doc.add_paragraph()
    run = body2_heading.add_run("Mobile Money as a Catalyst for Inclusion")
    run.bold = True
    run.font.size = Pt(12)

    body2_text = doc.add_paragraph(
        "The emergence of mobile money services has proven transformative for financial inclusion across the "
        "region. Kenya's M-Pesa, launched in 2007, demonstrated that mobile-based financial services could "
        "reach populations previously excluded from the formal financial system (4). Research has consistently "
        "shown that mobile money adoption correlates with increased household savings, improved ability to "
        "manage financial shocks, and higher levels of entrepreneurial activity. Ghana, Tanzania, and Uganda "
        "have similarly experienced dramatic expansions in financial access through mobile money platforms, "
        "with interoperability initiatives further enhancing utility across different providers. The low cost "
        "of mobile phone ownership relative to formal banking requirements has made these platforms accessible "
        "to previously excluded segments, including rural farmers, informal sector workers, and women "
        "entrepreneurs who face additional barriers to formal banking (5)."
    )
    body2_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- Body Paragraph 3 (contains '(ref needed)' placeholder) ---
    body3_heading = doc.add_paragraph()
    run = body3_heading.add_run("Fintech Innovation and Systemic Barriers")
    run.bold = True
    run.font.size = Pt(12)

    body3_text = doc.add_paragraph(
        "Despite the significant progress enabled by mobile money, deeper structural barriers continue to "
        "limit the transformative potential of fintech across sub-Saharan Africa. Regulatory fragmentation "
        "across national jurisdictions creates compliance burdens that deter cross-border fintech operations "
        "and limit the development of regional payment systems. Digital literacy deficits, particularly "
        "among older populations and those with limited formal education, restrict the uptake of more "
        "sophisticated financial products such as digital insurance and investment platforms (ref needed). "
        "Furthermore, persistent gender gaps in mobile phone ownership — with women on average 15% less "
        "likely than men to own a mobile phone in the region — continue to reproduce patterns of financial "
        "exclusion along gender lines. Infrastructure limitations, including unreliable electricity supply "
        "and patchy cellular coverage in remote areas, further constrain the reach of digital financial "
        "services (6)."
    )
    body3_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- Body Paragraph 4 ---
    body4_heading = doc.add_paragraph()
    run = body4_heading.add_run("Policy Implications and Future Directions")
    run.bold = True
    run.font.size = Pt(12)

    body4_text = doc.add_paragraph(
        "Addressing the remaining challenges requires coordinated policy action at both national and "
        "regional levels. Regulatory sandboxes have emerged as a promising mechanism for encouraging "
        "innovation while maintaining appropriate consumer protections, with Rwanda, Sierra Leone, and "
        "Mozambique among the early adopters of this approach (7). Regional harmonization of financial "
        "regulations, particularly through frameworks such as the African Continental Free Trade Area "
        "(AfCFTA), could significantly reduce compliance costs and enable the development of pan-African "
        "fintech solutions. Investments in digital literacy programs, particularly targeting women and "
        "rural populations, represent another priority for policymakers seeking to maximize the "
        "inclusive potential of fintech adoption. Finally, public-private partnerships to extend "
        "telecommunications infrastructure to underserved areas remain essential for ensuring that "
        "the benefits of fintech innovation are broadly shared."
    )
    body4_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- Conclusion ---
    conclusion_heading = doc.add_paragraph()
    run = conclusion_heading.add_run("Conclusion")
    run.bold = True
    run.font.size = Pt(12)

    conclusion_text = doc.add_paragraph(
        "Fintech has demonstrated remarkable potential to advance financial inclusion across sub-Saharan "
        "Africa, as evidenced by the region's position as a global leader in mobile money adoption. "
        "However, realizing the full transformative potential of these technologies requires sustained "
        "attention to structural barriers including regulatory complexity, digital literacy deficits, "
        "gender inequalities in technology access, and infrastructure gaps. The evidence reviewed in "
        "this essay suggests that coordinated policy interventions, combined with continued private "
        "sector innovation, offer the most promising pathway toward a financially inclusive future "
        "for the region."
    )
    conclusion_text.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- References Section ---
    ref_heading = doc.add_paragraph()
    run = ref_heading.add_run("References")
    run.bold = True
    run.font.size = Pt(12)

    references = [
        "1. Demirguc-Kunt, A., Klapper, L., Singer, D., Ansar, S., & Hess, J. (2022). The Global Findex "
        "Database 2021: Financial inclusion, digital payments, and resilience in the age of COVID-19. "
        "World Bank Publications. https://doi.org/10.1596/978-1-4648-1897-4",

        "2. Beck, T., Cull, R., & Jerome, A. (2021). Bank privatization and performance: Empirical "
        "evidence from Nigeria. Journal of Banking & Finance, 29(8-9), 2355-2379. "
        "https://doi.org/10.1016/j.jbankfin.2005.03.018",

        "3. Claessens, S., & Rojas-Suarez, L. (2020). Financial access and stability: Lessons from "
        "sub-Saharan Africa. Center for Global Development Working Paper, 310. "
        "https://www.cgdev.org/publication/financial-access-stability",

        "4. Jack, W., & Suri, T. (2011). Mobile money: The economics of M-PESA. NBER Working Paper "
        "No. 16721. National Bureau of Economic Research. https://doi.org/10.3386/w16721",

        "5. Suri, T., & Jack, W. (2016). The long-run poverty and gender impacts of mobile money. "
        "Science, 354(6317), 1288-1292. https://doi.org/10.1126/science.aah5309",

        "6. GSMA Intelligence. (2022). The mobile gender gap report 2022. GSMA. "
        "https://www.gsma.com/r/gender-gap/",

        "7. Ozili, P. K. (2022). Theories of financial inclusion. In Uncertainty and Challenges "
        "in Contemporary Economic Behaviour (pp. 89-115). Emerald Publishing Limited. "
        "https://doi.org/10.1108/978-1-80043-095-220211007",
    ]

    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
