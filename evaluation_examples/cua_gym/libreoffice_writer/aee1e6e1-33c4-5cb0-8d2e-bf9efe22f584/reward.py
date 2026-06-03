"""
Reward Script: Distribute all rows evenly (equal heights) and all columns evenly (equal widths)
Task ID: writer_tbl_022
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): All 4 rows have equal height via trHeight XML attribute
  Component 2 (0.4): All 4 columns have equal width (all cell tcW values equal)
  Precondition gate: Cell contents must be preserved (not scored independently)
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_022'

# Expected cell contents (from task context) — used as precondition gate only
EXPECTED_CELLS = [
    ['Category', 'Item', 'Details', 'Status'],
    ['Electronics', 'Laptop', 'High-performance laptop with 16GB RAM and 512GB SSD storage, suitable for professional use', 'Available'],
    ['Furniture', 'Chair', 'Office chair', 'Available'],
    ['Books', 'Novel', 'Fiction', 'Sold Out'],
]


def verify_task(file_path):
    """
    Verify that the table in the docx has equal row heights and equal column widths.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: file must have exactly one table with 4 rows and 4 columns
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    if len(table.rows) != 4 or len(table.columns) != 4:
        print(f"CRITICAL: Expected 4x4 table, found {len(table.rows)}x{len(table.columns)}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: Cell contents must be preserved (not scored independently,
    # as cell contents are the same in both initial and golden states)
    contents_ok = True
    try:
        for ridx, row in enumerate(table.rows):
            for cidx, cell in enumerate(row.cells):
                actual_text = cell.text.strip()
                expected_text = EXPECTED_CELLS[ridx][cidx].strip()
                if actual_text != expected_text:
                    print(f"GATE_FAIL: Cell [{ridx}][{cidx}] expected '{expected_text}', found '{actual_text}'")
                    contents_ok = False
        if contents_ok:
            print("GATE_PASS: All cell contents preserved (precondition)")
        else:
            print("GATE_FAIL: Cell contents corrupted — scoring continues but data integrity violated")
    except Exception as e:
        print(f"GATE_ERROR: Content check failed: {e}")

    # Component 1: All rows have equal height (0.6 points)
    # Task requires "distribute rows evenly so they have equal heights"
    # In the golden file, this is done by setting trHeight for all rows to the same value.
    # In the initial file, rows have no explicit trHeight set (uneven auto heights).
    try:
        row_heights = []
        for ridx, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is not None:
                trHeight = trPr.find(qn('w:trHeight'))
                if trHeight is not None:
                    val = trHeight.get(qn('w:val'))
                    row_heights.append(int(val) if val is not None else None)
                else:
                    row_heights.append(None)
            else:
                row_heights.append(None)

        # All row heights must be non-None (height is explicitly set)
        all_set = all(h is not None for h in row_heights)
        # All row heights must be equal
        non_none = [h for h in row_heights if h is not None]
        all_equal = len(set(non_none)) == 1 if non_none else False

        if all_set and all_equal:
            print(f"PASS: Component 1 — All 4 rows have equal height ({row_heights[0]} twips each) (0.6 pts)")
            total_score += 0.6
        elif not all_set:
            unset = [i for i, h in enumerate(row_heights) if h is None]
            print(f"FAIL: Component 1 — Rows {unset} have no explicit height set; heights={row_heights}")
        else:
            print(f"FAIL: Component 1 — Row heights are set but unequal: {row_heights}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All columns have equal width (0.4 points)
    # Task requires "distribute columns evenly so they have equal widths"
    # In the golden file, all cell tcW values are set to the same value (2340 twips).
    # In the initial file, cell widths are uneven: 1200, 1100, 5400, 1260 twips.
    try:
        all_cell_widths = []
        for ridx, row in enumerate(table.rows):
            row_widths = []
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        w_val = tcW.get(qn('w:w'))
                        row_widths.append(int(w_val) if w_val is not None else None)
                    else:
                        row_widths.append(None)
                else:
                    row_widths.append(None)
            all_cell_widths.append(row_widths)

        # Flatten all cell widths
        flat_widths = [w for row_w in all_cell_widths for w in row_w]
        all_set_widths = all(w is not None for w in flat_widths)
        # All widths must be equal
        non_none_widths = [w for w in flat_widths if w is not None]
        unique_widths = set(non_none_widths)
        all_equal_widths = len(unique_widths) == 1 if non_none_widths else False

        if all_set_widths and all_equal_widths:
            col_width_val = flat_widths[0]
            print(f"PASS: Component 2 — All 16 cells have equal width ({col_width_val} twips) (0.4 pts)")
            total_score += 0.4
        elif not all_set_widths:
            print(f"FAIL: Component 2 — Some cells have no explicit width set; widths={all_cell_widths}")
        else:
            print(f"FAIL: Component 2 — Cell widths are not all equal; unique values={unique_widths}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Run verification against the canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
