"""
Initial Setup: Distribute table rows and columns evenly
Task ID: writer_tbl_022
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_022'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_col_width(cell, width_twips):
    """Set a table cell's width in twips (twentieths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing tcW
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_row_height(row, height_twips, rule='auto'):
    """Set a table row's height in twips."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    # Remove any existing trHeight
    for existing in trPr.findall(qn('w:trHeight')):
        trPr.remove(existing)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), rule)
    trPr.append(trHeight)


def create_initial():
    doc = Document()

    # Set page margins for reference
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add a title paragraph
    title = doc.add_paragraph('Product Inventory')
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Add the table: 4 rows x 4 columns
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # Define cell content
    cell_data = [
        ['Category', 'Item', 'Details', 'Status'],
        ['Electronics', 'Laptop',
         'High-performance laptop with 16GB RAM and 512GB SSD storage, suitable for professional use',
         'Available'],
        ['Furniture', 'Chair', 'Office chair', 'Available'],
        ['Books', 'Novel', 'Fiction', 'Sold Out'],
    ]

    for r_idx, row_data in enumerate(cell_data):
        row = table.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = text
            if r_idx == 0:
                # Bold header row
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    # Set UNEVEN column widths (in twips; 1 inch = 1440 twips)
    # Total usable width ~6.5 inches = 9360 twips, intentionally uneven:
    # Col 0: 1200 twips (~0.83"), Col 1: 1100 twips (~0.76")
    # Col 2: 5400 twips (~3.75"), Col 3: 1260 twips (~0.875")
    uneven_widths = [1200, 1100, 5400, 1260]
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            set_col_width(cell, uneven_widths[c_idx])

    # Do NOT set fixed row heights — let them vary automatically with content
    # (Row 2 will be taller due to long text in Details cell)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
