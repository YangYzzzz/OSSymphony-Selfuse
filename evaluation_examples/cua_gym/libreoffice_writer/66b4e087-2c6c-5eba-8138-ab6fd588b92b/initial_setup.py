"""
Initial Setup: Apply alternating row background colors to a data table
Task ID: writer_tm_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_shading(cell, hex_color):
    """Set background/shading color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Quarterly Sales Report - Q1 2025', level=1)

    # Add introductory paragraph
    doc.add_paragraph(
        'The following table summarizes product sales performance across '
        'all regions for the first quarter of 2025. Please review the data '
        'and apply formatting as needed for the executive presentation.'
    )

    # Create 4-column, 12-row table (1 header + 11 data rows)
    table = doc.add_table(rows=12, cols=4)
    table.style = 'Table Grid'

    # Header row (Row 1) - blue background with white bold text
    headers = ['Product', 'Region', 'Units Sold', 'Revenue ($)']
    for j, header_text in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        # Set blue background
        set_cell_shading(cell, '4472C4')

    # Data rows (Rows 2-12) - all white background
    data = [
        ['Alpine Pro Jacket',     'Northeast',  '342',   '45,230.00'],
        ['TrailBlazer Boots',     'Southwest',   '287',   '31,570.00'],
        ['Summit Backpack',       'Pacific',     '519',   '38,925.00'],
        ['River Guide Kayak',     'Midwest',     '156',   '62,400.00'],
        ['Peak Performance Tent', 'Southeast',   '203',   '50,750.00'],
        ['Horizon Sunglasses',    'Northeast',   '891',   '22,275.00'],
        ['Arctic Sleeping Bag',   'Northwest',   '178',   '26,700.00'],
        ['Cascade Rain Shell',    'Pacific',     '445',   '53,400.00'],
        ['Expedition GPS Watch',  'Southwest',   '312',   '74,880.00'],
        ['Basecamp Cooler',       'Midwest',     '267',   '18,690.00'],
        ['Velocity Cycling Kit',  'Southeast',   '134',   '20,100.00'],
    ]

    for i, row_data in enumerate(data):
        row_idx = i + 1  # rows 1-11 in 0-indexed (rows 2-12 in task terms)
        for j, val in enumerate(row_data):
            cell = table.cell(row_idx, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            # Set white background explicitly for all data rows
            set_cell_shading(cell, 'FFFFFF')

    # Add a closing paragraph
    doc.add_paragraph('')
    doc.add_paragraph(
        'Note: Revenue figures are rounded to the nearest dollar. '
        'Regional breakdowns are available in the appendix.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
