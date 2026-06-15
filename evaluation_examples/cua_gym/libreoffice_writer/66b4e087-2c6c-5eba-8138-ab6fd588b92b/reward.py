"""
Reward Script: Apply alternating row background colors to data table
Task ID: writer_tm_024
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Even data rows (1-indexed 2,4,6,8,10,12) have #F0F0F0 background
  Component 2 (0.3): Odd data rows (1-indexed 3,5,7,9,11) have #FFFFFF background
  Component 3 (0.2): Header row (row 1) retains blue background (#4472C4)
  Component 4 (0.1): Cell content unchanged (spot-check)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_024'


def get_row_shading(row):
    """Get the shading fill color for all cells in a row. Returns list of fill strings."""
    fills = []
    for cell in row.cells:
        tc = cell._tc
        tc_pr = tc.find(qn('w:tcPr'))
        if tc_pr is not None:
            shd = tc_pr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                fills.append(fill.upper() if fill else None)
            else:
                fills.append(None)
        else:
            fills.append(None)
    return fills


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one table with 12 rows
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) < 12:
        print(f"FAIL: Expected 12 rows, found {len(table.rows)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Even data rows (1-indexed rows 2,4,6,8,10,12 = 0-indexed 1,3,5,7,9,11)
    # should have #F0F0F0 background (0.4 points)
    try:
        even_rows_idx = [1, 3, 5, 7, 9, 11]  # 0-indexed
        even_correct = 0
        even_total = len(even_rows_idx)
        for idx in even_rows_idx:
            fills = get_row_shading(table.rows[idx])
            if all(f == 'F0F0F0' for f in fills):
                even_correct += 1
            else:
                print(f"FAIL: Row {idx} (1-indexed {idx+1}) expected all F0F0F0, got {fills}")

        if even_correct == even_total:
            print(f"PASS: Component 1 — All {even_total} even data rows have #F0F0F0 background (0.4 pts)")
            total_score += 0.4
        elif even_correct > 0:
            partial = 0.4 * (even_correct / even_total)
            print(f"PARTIAL: Component 1 — {even_correct}/{even_total} even data rows correct ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No even data rows have #F0F0F0 background")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Odd data rows remain #FFFFFF AND even data rows have #F0F0F0 (0.3 points)
    # This is a compound check: odd rows white is only meaningful if even rows changed to gray
    # (prevents scoring on initial state where ALL data rows are white)
    try:
        odd_rows_idx = [2, 4, 6, 8, 10]  # 0-indexed
        odd_correct = 0
        odd_total = len(odd_rows_idx)
        for idx in odd_rows_idx:
            fills = get_row_shading(table.rows[idx])
            if all(f == 'FFFFFF' for f in fills):
                odd_correct += 1
            else:
                print(f"FAIL: Row {idx} (1-indexed {idx+1}) expected all FFFFFF, got {fills}")

        # Anchor: at least one even row must have changed to F0F0F0
        has_alternating = any(
            all(f == 'F0F0F0' for f in get_row_shading(table.rows[idx]))
            for idx in [1, 3, 5, 7, 9, 11]
        )

        if odd_correct == odd_total and has_alternating:
            print(f"PASS: Component 2 — All {odd_total} odd data rows have #FFFFFF with alternating pattern (0.3 pts)")
            total_score += 0.3
        elif odd_correct > 0 and has_alternating:
            partial = 0.3 * (odd_correct / odd_total)
            print(f"PARTIAL: Component 2 — {odd_correct}/{odd_total} odd data rows correct ({partial:.2f} pts)")
            total_score += partial
        elif not has_alternating:
            print(f"FAIL: Component 2 — No alternating pattern detected (even rows not changed to F0F0F0)")
        else:
            print(f"FAIL: Component 2 — No odd data rows have #FFFFFF background")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row retains blue background (0.2 points)
    # This is a compound check: header must be blue AND at least one even row must be F0F0F0
    # (If header is blue in both initial and golden, we anchor to the alternating pattern change)
    try:
        header_fills = get_row_shading(table.rows[0])
        header_blue = all(f == '4472C4' for f in header_fills)
        # This component only awards points if alternating pattern exists (even rows changed)
        has_alternating = any(
            all(f == 'F0F0F0' for f in get_row_shading(table.rows[idx]))
            for idx in [1, 3, 5, 7, 9, 11]
        )
        if header_blue and has_alternating:
            print(f"PASS: Component 3 — Header retains blue (#4472C4) with alternating pattern present (0.2 pts)")
            total_score += 0.2
        elif not header_blue:
            print(f"FAIL: Component 3 — Header background changed from blue, got {header_fills}")
        else:
            print(f"FAIL: Component 3 — Header is blue but no alternating pattern detected")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Cell content unchanged - spot check known values (0.1 points)
    # Only awards if alternating colors are present (anchored to task change)
    try:
        expected_cells = {
            (0, 0): 'Product',
            (1, 0): 'Alpine Pro Jacket',
            (5, 2): '203',
            (11, 3): '20,100.00',
        }
        content_mismatches = 0
        for (r, c), expected_text in expected_cells.items():
            actual = table.rows[r].cells[c].text.strip()
            if actual != expected_text:
                print(f"FAIL: Cell ({r},{c}) expected '{expected_text}', got '{actual}'")
                content_mismatches += 1

        # Only score if alternating pattern is present (task-introduced change)
        if content_mismatches == 0 and has_alternating:
            print(f"PASS: Component 4 — Cell content intact with alternating pattern present (0.1 pts)")
            total_score += 0.1
        elif content_mismatches > 0:
            print(f"FAIL: Component 4 — Cell content has been modified")
        else:
            print(f"FAIL: Component 4 — Content ok but no alternating pattern (pre-task state)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
