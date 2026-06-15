"""
Reward Script: Create weekly schedule table with formatted headers
Task ID: writer_tbl_035
Domain: libreoffice_writer
Scoring:
  - Component 1: Table exists with 6 rows x 7 columns (0.30 pts)
  - Component 2: Header row contains correct 7 column labels (0.20 pts)
  - Component 3: Time slots in column 1 rows 2-6 (0.20 pts)
  - Component 4: Header background is blue (0.15 pts)
  - Component 5: Header text is white and bold (0.15 pts)
"""

import os
from math import sqrt

# python-docx
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_035'
FILE_PATH = f'{WORKDIR}/weekly_planner.docx'


def color_distance(rgb_hex, target_rgb):
    """Compute Euclidean distance between a hex color string and a target (r,g,b) tuple."""
    r = int(rgb_hex[0:2], 16)
    g = int(rgb_hex[2:4], 16)
    b = int(rgb_hex[4:6], 16)
    return sqrt((r - target_rgb[0]) ** 2 + (g - target_rgb[1]) ** 2 + (b - target_rgb[2]) ** 2)


def get_cell_shading_fill(cell):
    """Return the fill hex color from cell shading, or None."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    if fill and fill.upper() not in ('AUTO', 'NONE', ''):
        return fill.upper()
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: file must be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document — task not started")
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # -----------------------------------------------------------------------
    # Component 1: Table dimensions — 6 rows x 7 columns (0.30 points)
    # -----------------------------------------------------------------------
    try:
        if num_rows == 6 and num_cols == 7:
            print(f"PASS: Component 1 — Table has correct dimensions 6 rows x 7 cols (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — Expected 6x7 table, found {num_rows}x{num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Header row contains correct 7 column labels (0.20 points)
    # -----------------------------------------------------------------------
    expected_headers = ['Time', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    try:
        actual_headers = [cell.text.strip() for cell in table.rows[0].cells]
        if actual_headers == expected_headers:
            print(f"PASS: Component 2 — Header row matches expected labels {actual_headers} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — Expected headers {expected_headers}, found {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Time slots in column 1, rows 2-6 (0.20 points)
    # -----------------------------------------------------------------------
    expected_times = ['8:00', '9:00', '10:00', '11:00', '12:00']
    try:
        actual_times = [table.rows[ri].cells[0].text.strip() for ri in range(1, 6)]
        if actual_times == expected_times:
            print(f"PASS: Component 3 — Time slots match {actual_times} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — Expected times {expected_times}, found {actual_times}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Header row cells have blue background (0.15 points)
    # Blue target: roughly (0, 0, 128) to (100, 149, 237) region
    # Actual golden uses 4472C4 = (68, 114, 196) which is clearly blue:
    #   B channel dominant, RGB distance from pure blue (0,0,255) < 120
    # -----------------------------------------------------------------------
    try:
        blue_count = 0
        for ci in range(num_cols):
            cell = table.rows[0].cells[ci]
            fill = get_cell_shading_fill(cell)
            if fill and len(fill) == 6:
                r = int(fill[0:2], 16)
                g = int(fill[2:4], 16)
                b = int(fill[4:6], 16)
                # Blue: B channel dominant and reasonably saturated
                if b > 80 and b > r and b > g:
                    blue_count += 1

        if blue_count == num_cols:
            print(f"PASS: Component 4 — All {num_cols} header cells have blue background (0.15 pts)")
            total_score += 0.15
        elif blue_count > 0:
            print(f"PARTIAL: Component 4 — Only {blue_count}/{num_cols} header cells have blue background")
        else:
            print(f"FAIL: Component 4 — No header cells have blue background (expected blue fill)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    # Component 5: Header text is white and bold (0.15 points)
    # White = font_color RGB close to FFFFFF = (255, 255, 255)
    # Both bold=True AND white color must hold for ALL 7 header cells
    # -----------------------------------------------------------------------
    try:
        white_bold_count = 0
        for ci in range(num_cols):
            cell = table.rows[0].cells[ci]
            for para in cell.paragraphs:
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    is_bold = run.bold is True
                    # Check font color via XML (rgb[0..2] for RGBColor tuple)
                    try:
                        rgb = run.font.color.rgb
                        # RGBColor is a tuple (r, g, b); use subscript access
                        is_white = (rgb[0] >= 200 and rgb[1] >= 200 and rgb[2] >= 200)
                    except Exception:
                        # Fallback: check via XML color element
                        try:
                            rPr = run._element.find(qn('w:rPr'))
                            color_el = rPr.find(qn('w:color')) if rPr is not None else None
                            val = color_el.get(qn('w:val')) if color_el is not None else None
                            if val and val not in ('auto', 'AUTO'):
                                r = int(val[0:2], 16)
                                g = int(val[2:4], 16)
                                b = int(val[4:6], 16)
                                is_white = (r >= 200 and g >= 200 and b >= 200)
                            else:
                                is_white = False
                        except Exception:
                            is_white = False
                    if is_bold and is_white:
                        white_bold_count += 1

        if white_bold_count == num_cols:
            print(f"PASS: Component 5 — All {num_cols} header cells have white bold text (0.15 pts)")
            total_score += 0.15
        elif white_bold_count > 0:
            print(f"PARTIAL: Component 5 — Only {white_bold_count}/{num_cols} header cells have white bold text")
        else:
            print(f"FAIL: Component 5 — Header text is not white bold (found {white_bold_count}/{num_cols})")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical file path on the VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
