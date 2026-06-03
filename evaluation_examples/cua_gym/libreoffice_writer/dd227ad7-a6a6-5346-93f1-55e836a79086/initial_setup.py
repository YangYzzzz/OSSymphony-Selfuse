"""
Initial Setup: Weekly planner document with heading only (no table yet)
Task ID: writer_tbl_035
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_035'
OUTPUT = f'{WORKDIR}/weekly_planner.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure the Desktop directory exists on VM
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add heading "My Weekly Schedule" in bold
    heading_para = doc.add_paragraph()
    run = heading_para.add_run("My Weekly Schedule")
    run.bold = True
    run.font.size = Pt(16)

    # Save the document — no table, just the heading
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
