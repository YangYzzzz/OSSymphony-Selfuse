"""
Initial Setup: Format monetary values in quarterly financial report
Task ID: writer_txtfmt_058
Domain: libreoffice_writer

Creates a quarterly financial report document with six monetary values,
all in 12pt Calibri regular (no special font formatting on any amounts).
The task will require the agent to change the font of the monetary values
to Liberation Mono.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_058'
# Task context says file should be on ~/Desktop/
OUTPUT = f'{WORKDIR}/Desktop/quarterly_financials.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)

    doc = Document()

    # --- Heading ---
    heading = doc.add_heading('Quarterly Financial Summary', level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(16)

    # --- Introduction paragraph ---
    intro = doc.add_paragraph()
    run = intro.add_run(
        'This report provides an overview of the financial performance across all business '
        'divisions for the current quarter. All figures are presented in USD.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    # --- Q1 Revenue Section ---
    doc.add_paragraph()
    sec1_heading = doc.add_heading('Q1 Revenue Performance', level=2)
    for run in sec1_heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)

    p1 = doc.add_paragraph()
    r1a = p1.add_run('North America division reported total revenue of ')
    r1a.font.name = 'Calibri'
    r1a.font.size = Pt(12)
    r1b = p1.add_run('$1,250,000')
    r1b.font.name = 'Calibri'
    r1b.font.size = Pt(12)
    r1c = p1.add_run(
        ' for the first quarter, representing a 12% increase compared to the same period '
        'last year. Strong performance in enterprise software licensing was the primary driver.'
    )
    r1c.font.name = 'Calibri'
    r1c.font.size = Pt(12)

    p2 = doc.add_paragraph()
    r2a = p2.add_run('The Asia-Pacific region generated ')
    r2a.font.name = 'Calibri'
    r2a.font.size = Pt(12)
    r2b = p2.add_run('$890,500')
    r2b.font.name = 'Calibri'
    r2b.font.size = Pt(12)
    r2c = p2.add_run(
        ' in Q1, driven by new customer acquisitions in Japan and South Korea. '
        'Expansion into emerging markets continues to show positive momentum.'
    )
    r2c.font.name = 'Calibri'
    r2c.font.size = Pt(12)

    # --- Q2 Revenue Section ---
    doc.add_paragraph()
    sec2_heading = doc.add_heading('Q2 Revenue Performance', level=2)
    for run in sec2_heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)

    p3 = doc.add_paragraph()
    r3a = p3.add_run('Combined revenue for Q2 reached ')
    r3a.font.name = 'Calibri'
    r3a.font.size = Pt(12)
    r3b = p3.add_run('$2,140,000')
    r3b.font.name = 'Calibri'
    r3b.font.size = Pt(12)
    r3c = p3.add_run(
        ', exceeding the projected target of $1.9 million. The EMEA region contributed '
        'significantly, with major contract renewals from three Fortune 500 clients.'
    )
    r3c.font.name = 'Calibri'
    r3c.font.size = Pt(12)

    p4 = doc.add_paragraph()
    r4a = p4.add_run('Operating expenses for Q2 totaled ')
    r4a.font.name = 'Calibri'
    r4a.font.size = Pt(12)
    r4b = p4.add_run('$675,300')
    r4b.font.name = 'Calibri'
    r4b.font.size = Pt(12)
    r4c = p4.add_run(
        ', which includes personnel costs, infrastructure investment, and marketing spend. '
        'Cost optimization initiatives are expected to reduce this by 8% in Q3.'
    )
    r4c.font.name = 'Calibri'
    r4c.font.size = Pt(12)

    # --- Annual Projections Section ---
    doc.add_paragraph()
    sec3_heading = doc.add_heading('Annual Projections and Outlook', level=2)
    for run in sec3_heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)

    p5 = doc.add_paragraph()
    r5a = p5.add_run(
        'Based on current trajectory, the full-year revenue projection has been revised upward to '
    )
    r5a.font.name = 'Calibri'
    r5a.font.size = Pt(12)
    r5b = p5.add_run('$3,015,800')
    r5b.font.name = 'Calibri'
    r5b.font.size = Pt(12)
    r5c = p5.add_run(
        '. This revision reflects stronger than expected performance in SaaS subscriptions '
        'and professional services across all regions.'
    )
    r5c.font.name = 'Calibri'
    r5c.font.size = Pt(12)

    p6 = doc.add_paragraph()
    r6a = p6.add_run(
        'The board has approved a capital investment budget of '
    )
    r6a.font.name = 'Calibri'
    r6a.font.size = Pt(12)
    r6b = p6.add_run('$1,890,200')
    r6b.font.name = 'Calibri'
    r6b.font.size = Pt(12)
    r6c = p6.add_run(
        ' for the remainder of the fiscal year, allocated primarily toward R&D initiatives '
        'and expanding the sales team in high-growth markets.'
    )
    r6c.font.name = 'Calibri'
    r6c.font.size = Pt(12)

    # --- Closing paragraph ---
    doc.add_paragraph()
    closing = doc.add_paragraph()
    r_close = closing.add_run(
        'Management remains confident in achieving the revised annual targets. '
        'Detailed divisional breakdowns are available in the supplementary appendix.'
    )
    r_close.font.name = 'Calibri'
    r_close.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
