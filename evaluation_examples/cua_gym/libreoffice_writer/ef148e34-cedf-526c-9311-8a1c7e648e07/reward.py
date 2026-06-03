"""
Reward Script: Change font of six monetary values to Liberation Mono
Task ID: writer_txtfmt_058
Domain: libreoffice_writer
Scoring:
  Component 1: Each of the 6 monetary amounts has font 'Liberation Mono' (0.1 pts each = 0.6 total)
  Component 2: Exactly 6 runs in the document have Liberation Mono font — verifying
               the change was both complete and surgical (no extra runs changed) (0.4 pts)
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_058'

# The six monetary values that must be changed to Liberation Mono
MONETARY_VALUES = [
    '$1,250,000',
    '$890,500',
    '$2,140,000',
    '$675,300',
    '$3,015,800',
    '$1,890,200',
]

TARGET_FONT = 'Liberation Mono'
EXPECTED_LIBERATION_MONO_COUNT = 6


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Change the font of all 6 monetary values to Liberation Mono.
    Initial state: All text uses Calibri — no runs have Liberation Mono (scores 0.0).
    Golden state: The 6 monetary runs use Liberation Mono; all other text stays Calibri (scores 1.0).
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: Each of the 6 monetary values has Liberation Mono font ---
    # 0.1 pts per monetary amount correctly formatted = up to 0.6 pts
    # This FAILS on initial (all Calibri, no Liberation Mono) and awards points
    # per correctly changed amount in golden.

    monetary_found = {v: False for v in MONETARY_VALUES}
    monetary_font = {v: None for v in MONETARY_VALUES}

    try:
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text.strip()
                if text in MONETARY_VALUES:
                    monetary_found[text] = True
                    monetary_font[text] = run.font.name

        # Also check inside tables (defensive: task mentions paragraphs only)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text = run.text.strip()
                            if text in MONETARY_VALUES:
                                monetary_found[text] = True
                                monetary_font[text] = run.font.name

        monetary_pass_count = 0
        for amount in MONETARY_VALUES:
            if not monetary_found[amount]:
                print(f"FAIL: Monetary amount '{amount}' not found as a discrete run in document")
            elif monetary_font[amount] == TARGET_FONT:
                print(f"PASS: '{amount}' font is '{TARGET_FONT}' (0.1 pts)")
                monetary_pass_count += 1
            else:
                print(f"FAIL: '{amount}' font is '{monetary_font[amount]}', expected '{TARGET_FONT}'")

        if monetary_pass_count > 0:
            total_score += round(monetary_pass_count * 0.1, 6)
        component1_score = round(monetary_pass_count * 0.1, 1)
        print(f"Component 1 subtotal: {component1_score}/0.6")

    except Exception as e:
        print(f"ERROR: Component 1 (monetary font check) failed: {e}")

    # --- Component 2: Exactly 6 runs in the document use Liberation Mono (0.4 pts) ---
    # This verifies the change was both COMPLETE (all 6 done) and SURGICAL (nothing extra changed).
    # On initial_env: 0 runs have Liberation Mono → count != 6 → FAIL (0.0)
    # On golden_env: exactly 6 runs have Liberation Mono → count == 6 → PASS (0.4)

    try:
        liberation_mono_runs = []

        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name == TARGET_FONT and run.text.strip():
                    liberation_mono_runs.append(run.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.name == TARGET_FONT and run.text.strip():
                                liberation_mono_runs.append(run.text.strip())

        count = len(liberation_mono_runs)
        if count == EXPECTED_LIBERATION_MONO_COUNT:
            print(f"PASS: Exactly {EXPECTED_LIBERATION_MONO_COUNT} runs use Liberation Mono "
                  f"(complete and surgical change) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Expected exactly {EXPECTED_LIBERATION_MONO_COUNT} Liberation Mono runs, "
                  f"found {count}: {liberation_mono_runs[:10]}")

        print(f"Component 2 subtotal: {'0.4' if count == EXPECTED_LIBERATION_MONO_COUNT else '0.0'}/0.4")

    except Exception as e:
        print(f"ERROR: Component 2 (Liberation Mono run count) failed: {e}")

    final_score = min(round(total_score, 6), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on VM
file_path = f'{WORKDIR}/Desktop/quarterly_financials.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
