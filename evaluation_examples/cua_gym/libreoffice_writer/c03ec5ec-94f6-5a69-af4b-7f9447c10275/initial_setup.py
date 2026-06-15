"""
Initial Setup: email_template.docx without any page border
Task ID: writer_page_060
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_page_060'
OUTPUT = f'{WORKDIR}/Desktop/email_template.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # --- Page setup: A4, portrait, margins all 2.54cm ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Explicitly remove any pgBorders element that may exist in the default template
    sectPr = section._sectPr
    for pgBorders in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(pgBorders)

    # --- Email Template Content ---

    # Sender info
    sender_para = doc.add_paragraph()
    sender_para.paragraph_format.space_after = Pt(0)
    run = sender_para.add_run('From: Alexandra Rivera <a.rivera@meridiantech.com>')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(0)
    run = date_para.add_run('Date: March 12, 2025')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    to_para = doc.add_paragraph()
    to_para.paragraph_format.space_after = Pt(0)
    run = to_para.add_run('To: Benjamin Hartwell <b.hartwell@globalpartners.net>')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    subject_para = doc.add_paragraph()
    subject_para.paragraph_format.space_after = Pt(12)
    run = subject_para.add_run('Subject: Q1 2025 Partnership Review and Strategic Alignment Proposal')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = True

    # Salutation
    sal_para = doc.add_paragraph()
    sal_para.paragraph_format.space_after = Pt(6)
    run = sal_para.add_run('Dear Benjamin,')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Body paragraphs
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(6)
    body1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body1.add_run(
        'I hope this message finds you well. I am writing to follow up on our conversation '
        'from last month regarding the proposed expansion of our collaboration into the Southeast '
        'Asian market. As we move into Q1 2025, I believe it is an excellent time to review our '
        'joint objectives and align our strategies for the coming quarter.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(6)
    body2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body2.add_run(
        'Our analytics team has completed the initial market assessment report, which highlights '
        'several high-potential opportunities in Vietnam and Indonesia. The report indicates a '
        'projected market entry cost reduction of approximately 18% compared to our original '
        'estimates, primarily due to favorable exchange rates and our existing logistics partnerships.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(6)
    body3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body3.add_run(
        'I would like to propose a virtual meeting during the week of March 24th to present these '
        'findings in detail and discuss the next steps. Please let me know your availability and '
        'preferred time zone for scheduling purposes. Our executive team is available Monday through '
        'Thursday between 9:00 AM and 5:00 PM SGT.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    body4 = doc.add_paragraph()
    body4.paragraph_format.space_after = Pt(6)
    body4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body4.add_run(
        'In the meantime, I have attached the preliminary market assessment document and the '
        'proposed partnership framework for your review. Should you have any questions or require '
        'additional information prior to our meeting, please do not hesitate to reach out to me '
        'directly at the contact details below.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Closing
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.space_after = Pt(6)
    run = closing_para.add_run('Thank you for your continued partnership and trust in Meridian Technologies.')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    regards_para = doc.add_paragraph()
    regards_para.paragraph_format.space_after = Pt(0)
    run = regards_para.add_run('Warm regards,')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(0)
    run = name_para.add_run('Alexandra Rivera')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = True

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(0)
    run = title_para.add_run('Senior Partnership Manager')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    company_para = doc.add_paragraph()
    company_para.paragraph_format.space_after = Pt(0)
    run = company_para.add_run('Meridian Technologies Pte. Ltd.')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    phone_para = doc.add_paragraph()
    phone_para.paragraph_format.space_after = Pt(0)
    run = phone_para.add_run('Phone: +65 6123 4567  |  Mobile: +65 9876 5432')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    email_para = doc.add_paragraph()
    email_para.paragraph_format.space_after = Pt(0)
    run = email_para.add_run('Email: a.rivera@meridiantech.com  |  www.meridiantech.com')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # NOTE: No page border is set — this is the initial state before the task
    # Verify: sectPr has no pgBorders
    assert sectPr.find(qn('w:pgBorders')) is None, 'Initial file must NOT have pgBorders!'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
