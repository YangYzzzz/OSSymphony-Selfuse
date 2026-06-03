"""
Initial Setup: Academic document with default page styles for all chapters
Task ID: writer_acad_085
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_085'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section):
    """Add a centered page number field to the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # PAGE field code
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def add_header(section, text):
    """Add a header with text to the section."""
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = text
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in hp.runs:
        run.font.size = Pt(10)
        run.font.italic = True


def create_initial():
    doc = Document()

    # Default page style for all sections: standard margins
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- Section 1 (default): Title Page / Chapter 1 ---
    section0 = doc.sections[0]
    section0.top_margin = Cm(2.54)
    section0.bottom_margin = Cm(2.54)
    section0.left_margin = Cm(3.81)
    section0.right_margin = Cm(2.54)

    # Header and footer enabled for all sections (default style)
    add_header(section0, "Advances in Neural Network Architectures")
    add_page_number_footer(section0)

    # Title
    title = doc.add_heading('Advances in Neural Network Architectures', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(24)

    # Author info
    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_after = Pt(6)
    run = author.add_run('Dr. Elena Vasquez')
    run.font.size = Pt(14)

    affil = doc.add_paragraph()
    affil.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil.paragraph_format.space_after = Pt(6)
    run = affil.add_run('Department of Computer Science, Stanford University')
    run.font.size = Pt(11)
    run.font.italic = True

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_after = Pt(24)
    run = date_para.add_run('March 2026')
    run.font.size = Pt(11)

    # Abstract
    abstract_heading = doc.add_heading('Abstract', level=1)
    abstract_text = doc.add_paragraph(
        'This paper presents a comprehensive survey of recent advances in neural network '
        'architectures, with particular emphasis on transformer-based models and their '
        'applications in natural language processing, computer vision, and multimodal '
        'learning. We analyze the evolution from early convolutional architectures to '
        'modern attention-based designs, highlighting key innovations that have driven '
        'state-of-the-art performance across multiple benchmarks. Our analysis covers '
        'architectural modifications, training strategies, and scaling laws that govern '
        'the behavior of large-scale neural networks.'
    )
    abstract_text.paragraph_format.space_after = Pt(12)

    # --- Chapter 1: Introduction ---
    doc.add_page_break()
    ch1 = doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_paragraph(
        'The field of deep learning has undergone remarkable transformations over the past '
        'decade. Beginning with the breakthrough performance of AlexNet in the 2012 ImageNet '
        'competition, neural network architectures have evolved from relatively simple '
        'convolutional designs to highly sophisticated systems capable of processing multiple '
        'modalities simultaneously. This evolution has been driven by a combination of '
        'architectural innovations, increased computational resources, and the availability '
        'of large-scale training datasets.'
    )

    doc.add_paragraph(
        'The introduction of the transformer architecture by Vaswani et al. (2017) marked a '
        'paradigm shift in the field. The self-attention mechanism, which allows the model to '
        'weigh the importance of different input positions regardless of their distance, has '
        'proven to be remarkably versatile. Initially designed for machine translation, '
        'transformers have since been adapted for image recognition, protein structure '
        'prediction, and even music generation.'
    )

    doc.add_paragraph(
        'In this survey, we systematically examine the major architectural innovations that '
        'have shaped the current landscape of deep learning. We organize our analysis around '
        'three key themes: (1) the transition from convolution to attention, (2) the scaling '
        'properties of modern architectures, and (3) the emergence of multimodal designs '
        'that bridge different data modalities.'
    )

    # --- Section 2 / Chapter 2 ---
    # Add a new section (page break) for Chapter 2
    sec_break = doc.add_paragraph()
    sec_break.paragraph_format.page_break_before = True
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.top_margin = Cm(2.54)
    new_section.bottom_margin = Cm(2.54)
    new_section.left_margin = Cm(3.81)
    new_section.right_margin = Cm(2.54)
    add_header(new_section, "Advances in Neural Network Architectures")
    add_page_number_footer(new_section)

    ch2 = doc.add_heading('Chapter 2: Convolutional Neural Networks', level=1)

    doc.add_paragraph(
        'Convolutional neural networks (CNNs) have been the backbone of computer vision '
        'research for over a decade. The hierarchical feature extraction paradigm, where '
        'early layers capture low-level features such as edges and textures while deeper '
        'layers represent high-level semantic concepts, has proven exceptionally effective '
        'for image classification, object detection, and semantic segmentation tasks.'
    )

    doc.add_heading('2.1 Architectural Evolution', level=2)
    doc.add_paragraph(
        'The evolution of CNN architectures can be traced through several landmark designs. '
        'VGGNet demonstrated that depth matters, with its 16-19 layer configurations achieving '
        'significant improvements over shallower networks. GoogLeNet introduced the inception '
        'module, which processes input at multiple scales simultaneously. ResNet solved the '
        'degradation problem with skip connections, enabling networks with hundreds of layers.'
    )

    doc.add_heading('2.2 Modern Variants', level=2)
    doc.add_paragraph(
        'Recent CNN variants have focused on efficiency and performance. EfficientNet uses '
        'compound scaling to balance network width, depth, and resolution. ConvNeXt modernizes '
        'the standard ConvNet design by incorporating ideas from transformers, such as larger '
        'kernel sizes and layer normalization, achieving competitive performance with vision '
        'transformers while maintaining the simplicity of convolutional architectures.'
    )

    # --- Section 3 / Chapter 3 ---
    sec_break2 = doc.add_paragraph()
    sec_break2.paragraph_format.page_break_before = True
    new_section2 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section2.top_margin = Cm(2.54)
    new_section2.bottom_margin = Cm(2.54)
    new_section2.left_margin = Cm(3.81)
    new_section2.right_margin = Cm(2.54)
    add_header(new_section2, "Advances in Neural Network Architectures")
    add_page_number_footer(new_section2)

    ch3 = doc.add_heading('Chapter 3: Transformer Architectures', level=1)

    doc.add_paragraph(
        'The transformer architecture has fundamentally changed the landscape of deep learning. '
        'Unlike recurrent neural networks that process sequences element by element, transformers '
        'compute representations of all positions simultaneously through self-attention. This '
        'parallel processing capability not only enables more efficient training on modern '
        'hardware but also allows the model to capture long-range dependencies more effectively.'
    )

    doc.add_heading('3.1 Self-Attention Mechanism', level=2)
    doc.add_paragraph(
        'The core innovation of the transformer is the scaled dot-product attention mechanism. '
        'Given query, key, and value matrices derived from the input, the attention output is '
        'computed as a weighted sum of values, where the weights are determined by the '
        'compatibility between queries and keys. Multi-head attention extends this by learning '
        'multiple attention patterns in parallel, each capturing different aspects of the input '
        'relationships.'
    )

    doc.add_heading('3.2 Vision Transformers', level=2)
    doc.add_paragraph(
        'The Vision Transformer (ViT) demonstrated that pure transformer architectures can '
        'achieve state-of-the-art performance on image classification when trained with '
        'sufficient data. By dividing images into fixed-size patches and treating them as '
        'tokens, ViT applies the standard transformer encoder directly to image data. '
        'Subsequent work, including DeiT, Swin Transformer, and BEiT, has addressed '
        'data efficiency and introduced hierarchical designs more suitable for dense '
        'prediction tasks.'
    )

    doc.add_heading('3.3 Large Language Models', level=2)
    doc.add_paragraph(
        'The scaling of transformer language models has yielded remarkable emergent capabilities. '
        'Models like GPT-4, PaLM, and Claude demonstrate that increasing model size, training '
        'data, and compute leads to qualitative improvements in reasoning, coding, and creative '
        'tasks. The scaling laws discovered by Kaplan et al. provide empirical guidance for '
        'optimal resource allocation during training, suggesting predictable relationships '
        'between model performance and computational budget.'
    )

    # References section
    doc.add_page_break()
    doc.add_heading('References', level=1)

    refs = [
        'Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention Is All You Need. NeurIPS.',
        'He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning. CVPR.',
        'Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al. (2021). An Image is Worth 16x16 Words. ICLR.',
        'Brown, T., Mann, B., Ryder, N., et al. (2020). Language Models are Few-Shot Learners. NeurIPS.',
        'Liu, Z., Lin, Y., Cao, Y., et al. (2021). Swin Transformer. ICCV.',
        'Tan, M. & Le, Q. (2019). EfficientNet: Rethinking Model Scaling. ICML.',
        'Kaplan, J., McCandlish, S., Henighan, T., et al. (2020). Scaling Laws for Neural Language Models.',
        'Simonyan, K. & Zisserman, A. (2015). Very Deep Convolutional Networks. ICLR.',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f'[{i}] {ref}')
        p.paragraph_format.space_after = Pt(4)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
