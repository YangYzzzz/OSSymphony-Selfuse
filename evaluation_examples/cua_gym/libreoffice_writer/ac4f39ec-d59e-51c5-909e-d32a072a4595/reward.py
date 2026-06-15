"""
Reward Script: Create a custom page style 'Chapter Start' with 5cm top margin,
               no header, and footer with centered page number.
Task ID: writer_acad_085
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35) — At least one section has top margin ~5cm
  Component 2 (0.35) — Chapter-start sections have header disabled (empty)
  Component 3 (0.30) — Chapter-start sections retain footer with PAGE field
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_085'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = list(doc.sections)
    if len(sections) < 2:
        print("FAIL: Document has fewer than 2 sections; cannot have chapter-start sections")
        print("REWARD: 0.0")
        return 0.0

    # Identify "chapter start" sections: those with top margin approximately 5cm
    # 5 cm = 1800000 EMU.  Allow tolerance of +/- 0.15 cm (54000 EMU).
    TARGET_TOP_EMU = 1800000  # 5 cm
    TOLERANCE = 54000         # ~0.15 cm

    chapter_start_indices = []
    for i, s in enumerate(sections):
        top = s.top_margin
        if top is not None and abs(top - TARGET_TOP_EMU) <= TOLERANCE:
            chapter_start_indices.append(i)

    # Component 1: At least one section has top margin ~5 cm (0.35 points)
    # This distinguishes golden (sections 1,2 have 5cm) from initial (all 2.54cm).
    try:
        if len(chapter_start_indices) > 0:
            cm_vals = [round(sections[i].top_margin / 360000, 2) for i in chapter_start_indices]
            print(f"PASS: Component 1 — {len(chapter_start_indices)} section(s) with ~5cm top margin "
                  f"(indices {chapter_start_indices}, values {cm_vals} cm) (0.35 pts)")
            total_score += 0.35
        else:
            all_tops = [round(s.top_margin / 360000, 2) if s.top_margin else None for s in sections]
            print(f"FAIL: Component 1 — No section has top margin ~5cm. All tops: {all_tops}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Chapter-start sections have NO header content (0.35 points)
    # Initial state: all sections have header text 'Advances in Neural Network Architectures'
    # Golden state: chapter-start sections have empty header
    try:
        if len(chapter_start_indices) == 0:
            print("FAIL: Component 2 — No chapter-start sections to check header on")
        else:
            all_headers_empty = True
            for idx in chapter_start_indices:
                header = sections[idx].header
                header_text = ""
                if header.paragraphs:
                    header_text = "".join(p.text for p in header.paragraphs).strip()
                if header_text:
                    all_headers_empty = False
                    print(f"FAIL: Component 2 — Section {idx} header is not empty: {repr(header_text[:60])}")
                    break

            if all_headers_empty:
                print(f"PASS: Component 2 — All chapter-start sections have empty/no header (0.35 pts)")
                total_score += 0.35
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Chapter-start sections retain footer with PAGE number field (0.30 points)
    # Both initial and golden have footers, but we gate this on chapter_start_indices existing
    # (which fails on initial), so this component only awards points when Comp 1 passes too.
    try:
        if len(chapter_start_indices) == 0:
            print("FAIL: Component 3 — No chapter-start sections to check footer on")
        else:
            all_footers_ok = True
            for idx in chapter_start_indices:
                footer = sections[idx].footer
                has_page_field = False
                if footer.paragraphs:
                    for para in footer.paragraphs:
                        for run in para.runs:
                            run_xml = run._element.xml
                            if 'instrText' in run_xml and 'PAGE' in run_xml:
                                has_page_field = True
                                break
                        if has_page_field:
                            break

                if not has_page_field:
                    all_footers_ok = False
                    print(f"FAIL: Component 3 — Section {idx} footer missing PAGE field code")
                    break

            if all_footers_ok:
                # Also check alignment is centered
                footer_centered = True
                for idx in chapter_start_indices:
                    for para in sections[idx].footer.paragraphs:
                        for run in para.runs:
                            if 'instrText' in run._element.xml and 'PAGE' in run._element.xml:
                                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                                align = para.paragraph_format.alignment
                                if align is not None and align != WD_PARAGRAPH_ALIGNMENT.CENTER:
                                    footer_centered = False
                                break

                if footer_centered:
                    print(f"PASS: Component 3 — Chapter-start sections have centered footer with PAGE field (0.30 pts)")
                    total_score += 0.30
                else:
                    print(f"FAIL: Component 3 — Footer page number is not centered")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
