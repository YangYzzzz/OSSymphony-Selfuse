"""
Initial Setup: Scientific research paper with technical terms needing footnotes
Task ID: writer_rd_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_028'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('Cellular Processes and Energy Conversion in Biological Systems', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Dr. Elena Vasquez, Dr. James Thornton, Dr. Mei-Lin Wu')
    run.font.size = Pt(11)
    run.italic = True

    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affiliation.add_run('Department of Molecular Biology, Pacific Northwest Research Institute')
    run.font.size = Pt(10)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Published: March 15, 2025')
    run.font.size = Pt(10)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This paper provides a comprehensive overview of the fundamental cellular processes '
        'that govern energy conversion and protein synthesis in eukaryotic organisms. We examine '
        'the interconnected pathways through which cells harvest light energy, generate adenosine '
        'triphosphate (ATP), and translate genetic information into functional proteins. Our analysis '
        'draws upon recent advances in cryo-electron microscopy, single-molecule fluorescence imaging, '
        'and computational modeling to present an integrated view of cellular metabolism.'
    )
    doc.add_paragraph(
        'Keywords: cellular biology, energy metabolism, protein synthesis, eukaryotic cells, '
        'organelle function, biochemical pathways'
    )

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The study of cellular processes has been a cornerstone of modern biology since the '
        'development of advanced microscopy techniques in the mid-twentieth century. As our '
        'understanding of molecular biology has deepened, researchers have uncovered intricate '
        'networks of biochemical reactions that sustain life at the cellular level. These processes '
        'range from the capture and conversion of solar energy to the precise assembly of complex '
        'protein molecules from amino acid building blocks.'
    )
    doc.add_paragraph(
        'In recent decades, technological advances have revolutionized our ability to observe '
        'and measure cellular processes with unprecedented precision. Cryo-electron microscopy '
        'has revealed the three-dimensional structures of molecular machines at near-atomic '
        'resolution, while single-molecule techniques have allowed researchers to track individual '
        'biochemical events in real time. These developments have led to a more nuanced '
        'understanding of how cells coordinate their diverse metabolic activities.'
    )
    doc.add_paragraph(
        'This review aims to synthesize current knowledge about three fundamental cellular '
        'processes: the conversion of light energy into chemical energy, the oxidative generation '
        'of ATP, and the translation of messenger RNA into polypeptide chains. We examine each '
        'process in detail, highlighting recent discoveries and their implications for our '
        'understanding of cellular function.'
    )
    doc.add_paragraph(
        'The significance of these processes cannot be overstated. Every living organism on Earth '
        'depends on the efficient operation of these biochemical pathways. Disruptions to any of '
        'these processes can lead to severe pathological conditions, including metabolic disorders, '
        'neurodegenerative diseases, and cancer. Understanding these processes at the molecular '
        'level is therefore essential not only for basic science but also for the development of '
        'novel therapeutic strategies.'
    )

    # --- Section 2: Light Energy Conversion (page 2 area) ---
    doc.add_heading('2. Light Energy Conversion in Plant Cells', level=1)
    doc.add_paragraph(
        'The capture and conversion of solar energy is one of the most remarkable biochemical '
        'processes found in nature. This process, which occurs primarily in the chloroplasts of '
        'plant cells and certain algae, represents the primary entry point of energy into most '
        'biological ecosystems. The molecular machinery responsible for this conversion has been '
        'refined over billions of years of evolution, resulting in a system of extraordinary '
        'efficiency and elegance.'
    )
    doc.add_paragraph(
        'The light-dependent reactions take place in the thylakoid membranes of chloroplasts. '
        'These membranes contain two photosystems, designated Photosystem I (PSI) and Photosystem II '
        '(PSII), each consisting of hundreds of chlorophyll and accessory pigment molecules arranged '
        'in precise antenna complexes. When a photon of light strikes an antenna pigment, the '
        'absorbed energy is transferred through a series of resonance events until it reaches the '
        'reaction center, where charge separation occurs.'
    )
    # Page 2, paragraph 3 - contains 'photosynthesis'
    doc.add_paragraph(
        'The overall process of photosynthesis is traditionally divided into two stages: the '
        'light-dependent reactions, which produce ATP and NADPH, and the Calvin cycle, which uses '
        'these energy carriers to fix carbon dioxide into organic molecules. Recent research by '
        'Tanaka et al. (2024) has revealed that the regulation of photosynthesis is far more '
        'complex than previously appreciated, involving sophisticated feedback mechanisms that '
        'adjust the rate of electron transport in response to fluctuating light conditions and '
        'metabolic demands of the cell.'
    )
    doc.add_paragraph(
        'The Calvin cycle, also known as the light-independent reactions, takes place in the '
        'stroma of the chloroplast. The key enzyme in this pathway is ribulose-1,5-bisphosphate '
        'carboxylase/oxygenase (RuBisCO), which catalyzes the fixation of atmospheric CO2 into '
        'a three-carbon compound. Despite being the most abundant enzyme on Earth, RuBisCO is '
        'remarkably slow, processing only about three molecules of CO2 per second. This limitation '
        'has significant implications for agricultural productivity and has been a target of '
        'bioengineering efforts aimed at improving crop yields.'
    )
    doc.add_paragraph(
        'Recent studies using time-resolved spectroscopy have provided new insights into the '
        'dynamics of energy transfer within photosynthetic antenna complexes. These studies have '
        'shown that quantum coherence effects may play a role in optimizing energy transfer '
        'efficiency, a finding that has implications for the design of artificial solar energy '
        'harvesting systems. The discovery of quantum effects in biological systems has opened '
        'up a new field of research known as quantum biology.'
    )

    # --- Section 3: Cellular Respiration (page 3 area) ---
    doc.add_heading('3. Oxidative Phosphorylation and Energy Production', level=1)
    # Page 3, paragraph 1 - contains 'mitochondria'
    doc.add_paragraph(
        'While plants and certain microorganisms can capture energy directly from sunlight, all '
        'eukaryotic cells rely on the mitochondria for the majority of their ATP production through '
        'a process known as oxidative phosphorylation. These double-membraned organelles, often '
        'referred to as the powerhouses of the cell, contain their own circular DNA and are believed '
        'to have originated from an ancient endosymbiotic relationship between a primitive eukaryotic '
        'cell and an alpha-proteobacterium approximately two billion years ago.'
    )
    doc.add_paragraph(
        'The process of oxidative phosphorylation occurs across the inner mitochondrial membrane, '
        'which is folded into numerous cristae to maximize surface area. The electron transport '
        'chain (ETC) consists of four major protein complexes (Complexes I through IV) and two '
        'mobile electron carriers (ubiquinone and cytochrome c). As electrons pass through the '
        'chain, protons are pumped from the matrix into the intermembrane space, creating an '
        'electrochemical gradient known as the proton motive force.'
    )
    doc.add_paragraph(
        'ATP synthase, sometimes called Complex V, is the molecular machine that harnesses the '
        'proton motive force to drive the phosphorylation of ADP to ATP. This remarkable rotary '
        'enzyme consists of two major subcomplexes: the F0 portion embedded in the inner membrane, '
        'which forms a proton channel, and the F1 portion that protrudes into the matrix and '
        'contains the catalytic sites for ATP synthesis. The rotation of the central stalk, driven '
        'by proton flow through F0, induces conformational changes in F1 that promote ATP formation.'
    )
    doc.add_paragraph(
        'The efficiency of oxidative phosphorylation is influenced by numerous factors, including '
        'the availability of substrates, the integrity of the inner membrane, and the activity of '
        'uncoupling proteins. Recent research has highlighted the role of mitochondrial dynamics '
        'in maintaining cellular energy homeostasis. The continuous cycle of fission and fusion '
        'allows damaged components to be segregated and degraded through mitophagy, while healthy '
        'segments are incorporated into the functional network.'
    )
    doc.add_paragraph(
        'Dysfunction of the mitochondrial electron transport chain has been implicated in a wide '
        'range of pathological conditions. Mutations in genes encoding ETC components can lead '
        'to severe metabolic disorders such as Leigh syndrome, MELAS (mitochondrial encephalomyopathy, '
        'lactic acidosis, and stroke-like episodes), and Kearns-Sayre syndrome. Furthermore, '
        'mitochondrial dysfunction has been linked to neurodegenerative diseases including '
        "Parkinson's disease, Alzheimer's disease, and amyotrophic lateral sclerosis (ALS)."
    )
    doc.add_paragraph(
        'The study of mitochondrial metabolism has also revealed important connections between '
        'energy production and cellular signaling. Reactive oxygen species (ROS) generated as '
        'byproducts of the electron transport chain serve as signaling molecules that regulate '
        'gene expression, cell proliferation, and apoptosis. The balance between ROS production '
        'and antioxidant defense systems is critical for maintaining cellular health, and '
        'disruptions to this balance are associated with aging and age-related diseases.'
    )

    # --- Section 4: Protein Synthesis (page 4 area) ---
    doc.add_heading('4. Protein Synthesis and the Translation Machinery', level=1)
    doc.add_paragraph(
        'The synthesis of proteins from messenger RNA (mRNA) templates is a fundamental process '
        'that occurs in all living cells. This process, known as translation, is carried out by '
        'a complex molecular machine that reads the genetic code and assembles amino acids into '
        'polypeptide chains with remarkable speed and accuracy. The fidelity of this process is '
        'essential for cellular function, as errors in protein synthesis can lead to misfolded '
        'or non-functional proteins with potentially harmful consequences.'
    )
    # Page 4, paragraph 2 - contains 'ribosome'
    doc.add_paragraph(
        'The central component of the translation machinery is the ribosome, a large '
        'ribonucleoprotein complex composed of two subunits. In eukaryotic cells, the small '
        '40S subunit and the large 60S subunit combine to form the functional 80S complex during '
        'translation initiation. The ribosome contains three tRNA binding sites: the aminoacyl '
        '(A) site, the peptidyl (P) site, and the exit (E) site. During elongation, the ribosome '
        'moves along the mRNA template in the 5\' to 3\' direction, reading codons and catalyzing '
        'peptide bond formation between successive amino acids.'
    )
    doc.add_paragraph(
        'Translation initiation is a highly regulated process that involves the assembly of the '
        'pre-initiation complex (PIC) at the mRNA start codon. In eukaryotes, this process requires '
        'at least twelve eukaryotic initiation factors (eIFs) that coordinate the binding of the '
        'initiator methionyl-tRNA to the 40S subunit and the subsequent scanning of the mRNA for '
        'the AUG start codon. The regulation of translation initiation is a major mechanism by '
        'which cells control gene expression in response to environmental signals.'
    )
    doc.add_paragraph(
        'The elongation phase of translation is catalyzed by elongation factors (eEF1A and eEF2 in '
        'eukaryotes) and involves a cyclic process of aminoacyl-tRNA selection, peptide bond '
        'formation, and translocation. The peptidyl transferase activity of the large ribosomal '
        'subunit is carried out by ribosomal RNA rather than protein, making this one of the most '
        'important examples of a ribozyme. This discovery provided strong support for the RNA '
        'world hypothesis, which proposes that RNA-based catalysts predated protein enzymes in '
        'the evolution of life.'
    )
    doc.add_paragraph(
        'Quality control mechanisms operating during and after translation ensure that only '
        'correctly synthesized proteins are released into the cell. Co-translational folding, '
        'assisted by molecular chaperones such as the Hsp70 and Hsp90 families, helps nascent '
        'polypeptides achieve their native conformations. Proteins that fail to fold properly '
        'are targeted for degradation by the ubiquitin-proteasome system or by autophagy-mediated '
        'pathways. These quality control systems are essential for maintaining cellular proteostasis.'
    )

    # --- Section 5: Discussion ---
    doc.add_heading('5. Integration of Cellular Processes', level=1)
    doc.add_paragraph(
        'The three processes described in this review are deeply interconnected and mutually '
        'dependent. The products of light energy conversion provide the carbon skeletons and '
        'energy carriers needed for cellular respiration, which in turn generates the ATP required '
        'for protein synthesis and other energy-demanding cellular activities. This metabolic '
        'interdependence underscores the importance of viewing cellular function as an integrated '
        'system rather than a collection of isolated pathways.'
    )
    doc.add_paragraph(
        'Emerging research has revealed additional layers of complexity in the regulation of '
        'these processes. Retrograde signaling from organelles to the nucleus allows cells to '
        'coordinate nuclear gene expression with the functional state of their energy-producing '
        'compartments. Metabolic intermediates serve as signaling molecules that modulate gene '
        'expression through epigenetic mechanisms, creating feedback loops that fine-tune cellular '
        'metabolism in response to changing conditions.'
    )
    doc.add_paragraph(
        'The implications of this integrated view of cellular metabolism extend beyond basic '
        'science to practical applications in medicine, agriculture, and biotechnology. '
        'Understanding the molecular details of these processes provides opportunities for '
        'developing targeted therapies for metabolic diseases, engineering more efficient crop '
        'plants, and designing novel bioenergy systems inspired by natural photosynthetic machinery.'
    )

    # --- References ---
    doc.add_heading('References', level=1)
    references = [
        'Alberts, B., Johnson, A., Lewis, J., et al. (2022). Molecular Biology of the Cell, 7th ed. W.W. Norton.',
        'Blankenship, R.E. (2021). Molecular Mechanisms of Photosynthesis, 3rd ed. Wiley-Blackwell.',
        'Krebs, J.E., Goldstein, E.S., Kilpatrick, S.T. (2023). Lewin\'s Genes XIII. Jones & Bartlett Learning.',
        'Nicholls, D.G. & Ferguson, S.J. (2023). Bioenergetics, 5th ed. Academic Press.',
        'Ramakrishnan, V. (2023). "Ribosome structure and the mechanism of translation." Annual Review of Biochemistry, 92, 123-150.',
        'Tanaka, R., Ito, H., & Tanaka, A. (2024). "Regulation of photosynthetic electron transport under fluctuating light." Nature Plants, 10(2), 145-158.',
        'Voet, D. & Voet, J.G. (2024). Biochemistry, 6th ed. Wiley.',
        'Youle, R.J. & van der Bliek, A.M. (2022). "Mitochondrial fission, fusion, and stress." Science, 337(6098), 1062-1065.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
