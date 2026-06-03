"""
Initial Setup: Insert a rectangle shape on page 1 with gradient fill
Task ID: writer_obj_030
Domain: libreoffice_writer

Creates header_design.docx at ~/Desktop/ with a title and body text on page 1,
but NO shapes (the agent will add the rectangle shape).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'header_design'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Create Desktop directory if needed
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Add a title
    title_para = doc.add_heading('Annual Corporate Report 2025', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a subtitle
    subtitle = doc.add_paragraph('Strategic Vision & Performance Overview')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Add spacing
    doc.add_paragraph('')

    # Add body paragraphs
    intro = doc.add_paragraph(
        'This report provides a comprehensive overview of our company\'s performance '
        'throughout the fiscal year 2025. We have achieved significant milestones in '
        'product development, market expansion, and financial growth.'
    )
    intro.paragraph_format.space_after = Pt(8)

    # Section heading
    section_heading = doc.add_heading('Executive Summary', level=1)

    body1 = doc.add_paragraph(
        'Our strategic initiatives have yielded exceptional results this year. '
        'Revenue grew by 18% year-over-year, reaching a record high of $4.2 billion. '
        'Operating margins improved to 24%, driven by operational efficiencies and '
        'disciplined cost management across all business units.'
    )
    body1.paragraph_format.space_after = Pt(8)

    body2 = doc.add_paragraph(
        'Customer satisfaction scores reached an all-time high of 94%, reflecting '
        'our commitment to delivering exceptional value and service quality. '
        'Our team of 12,000 dedicated employees continues to be our greatest asset, '
        'with employee engagement scores climbing to 87% globally.'
    )
    body2.paragraph_format.space_after = Pt(8)

    # Another section
    section_heading2 = doc.add_heading('Key Achievements', level=1)

    body3 = doc.add_paragraph(
        'The launch of our next-generation product platform in Q2 2025 has been met '
        'with overwhelming market enthusiasm. Over 500,000 enterprise customers have '
        'already adopted the new platform, representing a 35% faster adoption rate '
        'compared to previous major releases.'
    )
    body3.paragraph_format.space_after = Pt(8)

    body4 = doc.add_paragraph(
        'International expansion efforts have resulted in entry into 12 new markets '
        'across Asia-Pacific and Latin America. These markets now contribute 22% of '
        'total revenue, up from 15% in the previous fiscal year, demonstrating the '
        'strength of our global growth strategy.'
    )
    body4.paragraph_format.space_after = Pt(8)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup - open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
