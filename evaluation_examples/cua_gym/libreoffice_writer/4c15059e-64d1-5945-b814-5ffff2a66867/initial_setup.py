"""
Initial Setup: Apply consistent paragraph spacing to Corporate_Policy.docx
Task ID: writer_pd_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Helper functions ---
    def add_heading1(text):
        p = doc.add_heading(text, level=1)
        # Inconsistent spacing for heading 1: random values
        pf = p.paragraph_format
        pf.space_before = Pt(30)  # wrong - should be 24
        pf.space_after = Pt(18)   # wrong - should be 12
        return p

    def add_heading2(text):
        p = doc.add_heading(text, level=2)
        pf = p.paragraph_format
        pf.space_before = Pt(10)  # wrong - should be 18
        pf.space_after = Pt(14)   # wrong - should be 8
        return p

    def add_body(text):
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        # Inconsistent: left aligned, no indent, random spacing
        pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        pf.first_line_indent = None
        # Random spacing - sometimes too much, sometimes too little
        import random
        random.seed(hash(text) % 10000)
        pf.space_before = Pt(random.choice([0, 3, 6, 10, 12]))
        pf.space_after = Pt(random.choice([0, 3, 8, 10, 14]))
        return p

    def add_list_item(text):
        p = doc.add_paragraph(text, style='List Bullet')
        pf = p.paragraph_format
        # Inconsistent list spacing
        import random
        random.seed(hash(text) % 10000)
        pf.space_before = Pt(random.choice([0, 4, 6, 8]))
        pf.space_after = Pt(random.choice([0, 2, 6, 8]))
        return p

    def add_blank_line():
        """Add extra blank paragraph to simulate messy formatting."""
        p = doc.add_paragraph('')
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        return p

    # ======== DOCUMENT CONTENT (16 pages of corporate policy) ========

    # Title page
    title = doc.add_heading('Meridian Global Solutions', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(24)

    subtitle = doc.add_paragraph('Corporate Policy Manual')
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    version_info = doc.add_paragraph('Version 4.2 | Effective: January 15, 2025')
    version_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_info.paragraph_format.space_after = Pt(36)

    add_body('This document outlines the corporate policies, procedures, and guidelines applicable to all employees of Meridian Global Solutions. All personnel are expected to familiarize themselves with the contents of this manual and comply with the standards described herein.')

    add_blank_line()

    add_body('For questions regarding policy interpretation or application, please contact the Human Resources department at hr@meridianglobal.com or extension 4500.')

    doc.add_page_break()

    # ---- Section 1: Code of Conduct ----
    add_heading1('1. Code of Conduct')
    add_blank_line()
    add_body('Meridian Global Solutions is committed to maintaining the highest standards of ethical behavior in all business dealings. Every employee, contractor, and partner is expected to adhere to the principles outlined in this section.')

    add_heading2('1.1 Professional Behavior')
    add_body('All employees shall conduct themselves in a professional manner that reflects positively on the organization. This includes maintaining respectful communication with colleagues, clients, and external stakeholders at all times.')
    add_blank_line()
    add_body('Professionalism extends to digital communications including email, instant messaging, and video conferencing. Employees should ensure that all written and verbal communications are courteous, clear, and appropriate for a business environment.')

    add_heading2('1.2 Conflict of Interest')
    add_body('Employees must avoid situations where personal interests may conflict with the interests of the company. Any potential conflict of interest must be disclosed to the employee\'s direct supervisor and the Ethics Committee within five business days of identification.')

    add_list_item('Financial interests in competing companies exceeding $5,000 must be reported')
    add_list_item('Family relationships with vendors or suppliers require immediate disclosure')
    add_list_item('Outside employment must be approved by department management in advance')
    add_list_item('Board positions at external organizations require written authorization')
    add_blank_line()

    add_heading2('1.3 Anti-Harassment Policy')
    add_body('Meridian Global Solutions maintains a zero-tolerance policy toward harassment of any kind. This includes but is not limited to harassment based on race, gender, sexual orientation, religion, national origin, disability, or age.')
    add_body('Any employee who witnesses or experiences harassment should report the incident through one of the following channels:')

    add_list_item('Direct supervisor or department manager')
    add_list_item('Human Resources department (ext. 4500)')
    add_list_item('Anonymous ethics hotline: 1-888-555-ETHIC')
    add_list_item('Online reporting portal at ethics.meridianglobal.com')

    add_body('All reports will be investigated promptly and confidentially. Retaliation against individuals who report harassment in good faith is strictly prohibited and will result in disciplinary action up to and including termination.')

    doc.add_page_break()

    # ---- Section 2: Employment Policies ----
    add_heading1('2. Employment Policies')
    add_body('This section outlines the employment policies governing the relationship between Meridian Global Solutions and its workforce. These policies apply to all full-time and part-time employees across all departments and locations.')
    add_blank_line()

    add_heading2('2.1 Hiring and Onboarding')
    add_body('All hiring decisions are made based on qualifications, experience, and alignment with organizational values. The recruitment process follows a standardized framework to ensure fairness and consistency across all departments.')
    add_body('New employees undergo a comprehensive onboarding program lasting approximately four weeks. During this period, employees will complete the following milestones:')

    add_list_item('Complete all required HR documentation and benefits enrollment within the first three days')
    add_list_item('Attend mandatory orientation sessions covering company history, culture, and organizational structure')
    add_list_item('Complete department-specific training modules assigned by their direct supervisor')
    add_list_item('Meet with assigned mentor for initial goal-setting and role expectations discussion')
    add_list_item('Submit signed acknowledgment of Corporate Policy Manual receipt')
    add_blank_line()

    add_heading2('2.2 Work Hours and Attendance')
    add_body('Standard business hours are Monday through Friday, 8:30 AM to 5:30 PM local time, with a one-hour lunch break. Flexible work arrangements may be available with managerial approval and must be documented through the HR portal.')
    add_body('Employees are expected to maintain regular attendance patterns. Excessive absences, defined as more than six unexcused absences in a rolling 12-month period, may result in progressive disciplinary action.')
    add_blank_line()

    add_heading2('2.3 Remote Work Policy')
    add_body('Eligible employees may request remote work arrangements for up to three days per week. Remote work eligibility is determined by job function, performance history, and departmental needs. All remote workers must maintain a dedicated workspace with reliable internet connectivity of at least 50 Mbps download speed.')

    add_list_item('Remote work agreements must be renewed quarterly with supervisor approval')
    add_list_item('Employees must be available during core hours (10:00 AM to 3:00 PM local time)')
    add_list_item('All company data accessed remotely must use the approved VPN connection')
    add_list_item('Home office equipment stipend of $500 is available annually for approved remote workers')

    add_heading2('2.4 Performance Reviews')
    add_body('Performance evaluations are conducted semi-annually in June and December. The review process incorporates self-assessment, peer feedback, and managerial evaluation using the company\'s proprietary performance framework.')
    add_body('Performance ratings follow a five-point scale: Exceptional (5), Exceeds Expectations (4), Meets Expectations (3), Needs Improvement (2), and Unsatisfactory (1). Employees receiving a rating of 2 or below will be placed on a Performance Improvement Plan lasting 90 days.')

    doc.add_page_break()

    # ---- Section 3: Compensation and Benefits ----
    add_heading1('3. Compensation and Benefits')
    add_blank_line()
    add_body('Meridian Global Solutions offers a competitive compensation and benefits package designed to attract and retain top talent. The total rewards philosophy emphasizes both monetary and non-monetary components to support employee wellbeing and professional growth.')

    add_heading2('3.1 Salary Structure')
    add_body('Salaries are benchmarked annually against industry standards using data from multiple compensation surveys. The company targets the 60th percentile of the market for base compensation and the 75th percentile for total compensation including bonuses and equity.')
    add_body('Salary adjustments are made during the annual compensation review cycle in March. Merit increases are based on performance ratings, market movement, and internal equity considerations. The typical merit increase range is 2-8% based on performance rating.')
    add_blank_line()

    add_heading2('3.2 Benefits Overview')
    add_body('Full-time employees working 30 or more hours per week are eligible for the complete benefits package. Benefits enrollment occurs during the annual open enrollment period in November or within 30 days of a qualifying life event.')

    add_list_item('Medical insurance: PPO and HMO options with company covering 80% of premiums for employee-only coverage')
    add_list_item('Dental and vision insurance with comprehensive coverage options')
    add_list_item('401(k) retirement plan with company match up to 6% of eligible compensation')
    add_list_item('Life insurance at 2x annual salary provided at no cost to employees')
    add_list_item('Short-term and long-term disability insurance')
    add_list_item('Employee Assistance Program providing confidential counseling and support services')
    add_list_item('Tuition reimbursement up to $8,000 annually for approved degree programs')
    add_blank_line()

    add_heading2('3.3 Paid Time Off')
    add_body('The company provides a generous paid time off program that includes vacation days, sick leave, and personal days. Accrual rates increase with tenure as outlined below:')
    add_body('Years 0-2: 15 vacation days, 8 sick days, 3 personal days per year. Years 3-5: 20 vacation days, 10 sick days, 3 personal days. Years 6-10: 25 vacation days, 12 sick days, 4 personal days. Years 11+: 30 vacation days, 15 sick days, 5 personal days.')
    add_body('Unused vacation days may be carried over up to a maximum of 10 days. Sick leave does not carry over but unused sick days are converted to a wellness bonus at a rate of $100 per unused day at year end.')

    add_heading2('3.4 Bonuses and Incentives')
    add_body('Annual performance bonuses are awarded based on individual performance and company financial results. Target bonus percentages vary by level: Individual Contributors 10%, Managers 15%, Directors 20%, Vice Presidents 25%, and C-Suite 35% of base salary.')

    doc.add_page_break()

    # ---- Section 4: Information Security ----
    add_heading1('4. Information Security')
    add_body('Protecting company information and client data is the responsibility of every employee. This section outlines the security policies and procedures that must be followed to maintain the integrity and confidentiality of all organizational data assets.')
    add_blank_line()

    add_heading2('4.1 Data Classification')
    add_body('All company data is classified into four categories based on sensitivity and regulatory requirements. Employees must handle data according to its classification level at all times.')

    add_list_item('Public: Information approved for external distribution (marketing materials, press releases)')
    add_list_item('Internal: Information intended for employee use only (internal memos, org charts)')
    add_list_item('Confidential: Sensitive business information (financial reports, strategic plans, client contracts)')
    add_list_item('Restricted: Highly sensitive data requiring special handling (PII, PHI, trade secrets, source code)')

    add_heading2('4.2 Password and Access Requirements')
    add_body('All employees must follow the password policy mandated by the Information Security team. Passwords must be a minimum of 14 characters and include a combination of uppercase letters, lowercase letters, numbers, and special characters.')
    add_blank_line()
    add_body('Multi-factor authentication is required for all systems containing Confidential or Restricted data. Authentication tokens must not be shared between individuals under any circumstances.')

    add_list_item('Passwords must be changed every 90 days and cannot repeat any of the last 12 passwords used')
    add_list_item('Account lockout occurs after five consecutive failed login attempts')
    add_list_item('Screen lock must activate after five minutes of inactivity')
    add_list_item('Shared accounts are prohibited except where explicitly approved by the CISO')
    add_blank_line()

    add_heading2('4.3 Device Security')
    add_body('All company-issued devices must have full disk encryption enabled and approved endpoint protection software installed. Personal devices used for company work must be enrolled in the Mobile Device Management system before accessing any company resources.')
    add_body('Lost or stolen devices must be reported to IT Security within two hours of discovery. The IT Security team will initiate remote wipe procedures for any device containing company data that cannot be recovered within 24 hours.')

    add_heading2('4.4 Incident Response')
    add_body('Any suspected security incident must be reported immediately to the Security Operations Center at security@meridianglobal.com or by calling the 24/7 hotline at 1-888-555-SECURE. Employees should not attempt to investigate or remediate security incidents on their own.')
    add_body('The incident response process follows a five-phase approach: Identification, Containment, Eradication, Recovery, and Lessons Learned. All incidents are documented and reviewed by the Security Review Board on a monthly basis.')

    doc.add_page_break()

    # ---- Section 5: Travel and Expense ----
    add_heading1('5. Travel and Expense Policy')
    add_blank_line()
    add_body('Business travel and associated expenses must be pre-approved by the employee\'s direct supervisor and comply with the guidelines established in this section. The company reimburses reasonable business expenses incurred while conducting company business.')

    add_heading2('5.1 Travel Authorization')
    add_body('All business travel must be requested through the corporate travel management system at least 10 business days prior to departure. Domestic travel requires supervisor approval; international travel requires both supervisor and department head approval.')
    add_body('Travel arrangements should be made using preferred vendors wherever possible. The company has negotiated corporate rates with several airlines, hotel chains, and car rental agencies. Use of non-preferred vendors requires documented justification.')

    add_heading2('5.2 Expense Categories and Limits')
    add_body('The following daily expense limits apply to all employees regardless of level:')

    add_list_item('Airfare: Economy class for flights under 6 hours; business class permitted for flights over 6 hours with VP approval')
    add_list_item('Hotel: Up to $250 per night for domestic travel; international rates based on per diem schedule by country')
    add_list_item('Ground transportation: Actual cost for taxis, rideshare, or rental cars; receipts required for amounts over $25')
    add_list_item('Meals: Up to $75 per day; itemized receipts required; alcohol is not reimbursable')
    add_list_item('Incidentals: Up to $30 per day for tips, parking, and miscellaneous expenses')
    add_blank_line()

    add_heading2('5.3 Expense Reporting')
    add_body('Expense reports must be submitted within 15 business days of trip completion using the online expense management system. All expenses over $25 require original receipts or digital copies uploaded to the system.')
    add_body('Managers must review and approve expense reports within 5 business days of submission. Approved expenses are reimbursed via direct deposit within two pay periods following approval.')

    doc.add_page_break()

    # ---- Section 6: Workplace Safety ----
    add_heading1('6. Workplace Safety and Health')
    add_body('Meridian Global Solutions is committed to providing a safe and healthy work environment for all employees, contractors, and visitors. The company complies with all applicable federal, state, and local occupational safety regulations.')

    add_heading2('6.1 General Safety Guidelines')
    add_body('All employees are responsible for maintaining a safe workplace by following established safety procedures and reporting hazards promptly. Safety training is mandatory for all new employees and must be refreshed annually.')
    add_blank_line()
    add_body('Emergency exits must remain clear and unobstructed at all times. Fire extinguishers, first aid kits, and AED devices are located throughout all facilities and should not be blocked or moved from designated locations.')

    add_list_item('Report unsafe conditions to Facilities Management at ext. 3200 or safety@meridianglobal.com')
    add_list_item('Participate in quarterly fire drills and annual emergency preparedness exercises')
    add_list_item('Know the location of the nearest emergency exit, fire extinguisher, and first aid kit')
    add_list_item('Keep workspaces clean and free of tripping hazards')

    add_heading2('6.2 Ergonomics')
    add_body('The company provides ergonomic assessments for all workstation setups. Employees experiencing discomfort related to their workspace should request an ergonomic evaluation through the HR portal. Common adjustments include monitor height, chair settings, keyboard positioning, and lighting modifications.')
    add_body('Standing desk converters are available upon request and subject to availability. Employees with documented medical conditions may request specialized ergonomic equipment through the reasonable accommodation process.')
    add_blank_line()

    add_heading2('6.3 Emergency Procedures')
    add_body('In the event of an emergency, employees should follow the posted evacuation procedures for their building. Each floor has designated Emergency Wardens responsible for coordinating evacuations and conducting headcounts at assembly points.')
    add_body('The company maintains an Emergency Notification System that sends alerts via text message, email, and phone to all employees during critical events. Employees must keep their contact information current in the HR system to ensure timely notification.')

    doc.add_page_break()

    # ---- Section 7: Intellectual Property ----
    add_heading1('7. Intellectual Property')
    add_blank_line()
    add_body('All intellectual property created by employees during the course of their employment belongs to Meridian Global Solutions. This includes inventions, designs, software, documentation, processes, and any other creative works developed using company resources or related to company business.')

    add_heading2('7.1 Confidential Information')
    add_body('Employees are required to protect confidential information both during and after their employment with the company. The Non-Disclosure Agreement signed at the time of hiring remains in effect for a period of three years following separation from the company.')
    add_body('Confidential information includes but is not limited to: customer lists, pricing strategies, financial data, product roadmaps, source code, algorithms, business processes, and strategic plans. Employees must not disclose confidential information to any unauthorized person or entity.')
    add_blank_line()

    add_heading2('7.2 Invention Assignment')
    add_body('All inventions, improvements, and discoveries made by employees that relate to the company\'s current or anticipated business activities are assigned to the company. Employees must promptly disclose any such inventions to their supervisor and the Legal department.')

    add_list_item('Pre-existing intellectual property must be disclosed during the onboarding process')
    add_list_item('Open source contributions require prior approval from the Engineering Legal Liaison')
    add_list_item('Publication of research papers or articles related to company work requires Legal review')
    add_list_item('Social media posts about company technology must comply with the Social Media Policy')

    add_heading2('7.3 Trade Secrets')
    add_body('The company\'s trade secrets represent significant competitive advantages and must be protected with the highest level of care. Access to trade secret information is granted on a need-to-know basis and requires explicit authorization from the relevant department head.')
    add_body('Employees who handle trade secret information must take reasonable precautions to prevent unauthorized disclosure, including using encrypted communications, securing physical documents in locked cabinets, and limiting discussions of sensitive topics to secure environments.')

    doc.add_page_break()

    # ---- Section 8: Corporate Social Responsibility ----
    add_heading1('8. Corporate Social Responsibility')
    add_body('Meridian Global Solutions recognizes its obligation to operate in a socially responsible manner. The company is committed to environmental sustainability, community engagement, and ethical business practices across all operations and supply chains.')

    add_heading2('8.1 Environmental Policy')
    add_body('The company has established ambitious sustainability targets aligned with the Paris Agreement goals. By 2030, Meridian aims to reduce Scope 1 and Scope 2 greenhouse gas emissions by 50% from the 2020 baseline and achieve carbon neutrality for all global operations by 2040.')
    add_blank_line()
    add_body('All offices participate in the Green Office Initiative which includes recycling programs, energy conservation measures, and sustainable procurement practices. Employees are encouraged to minimize paper usage, utilize digital tools for collaboration, and report energy waste.')

    add_list_item('Single-use plastics have been eliminated from all company cafeterias and break rooms')
    add_list_item('Electric vehicle charging stations are available at all major office locations')
    add_list_item('The company purchases 100% renewable energy certificates for all domestic facilities')
    add_list_item('Business travel carbon offset program automatically applies to all booked flights')

    add_heading2('8.2 Community Engagement')
    add_body('The company encourages employees to participate in community service activities and provides 16 hours of paid volunteer time annually. The Meridian Foundation distributes over $2 million annually to nonprofit organizations focused on education, healthcare, and environmental conservation.')
    add_body('Employee matching gift program doubles personal charitable contributions up to $2,500 per calendar year. Team volunteer events are organized quarterly by the Community Engagement Committee.')
    add_blank_line()

    add_heading2('8.3 Diversity, Equity, and Inclusion')
    add_body('Meridian Global Solutions is committed to building a diverse and inclusive workplace where all individuals feel valued, respected, and empowered to contribute their best work. The DEI Council oversees strategic initiatives, tracks progress metrics, and reports to the Board of Directors quarterly.')
    add_body('The company has established the following DEI targets for 2026: Achieve 45% representation of women in leadership positions, increase underrepresented minority representation to 30% across all levels, and maintain a disability inclusion score above 80 on the Disability Equality Index.')

    doc.add_page_break()

    # ---- Section 9: Disciplinary Procedures ----
    add_heading1('9. Disciplinary Procedures')
    add_blank_line()
    add_body('The company follows a progressive discipline approach to address performance and behavioral issues. The goal of disciplinary action is to correct behavior and help employees succeed, not to punish. All disciplinary actions are documented and maintained in the employee\'s personnel file.')

    add_heading2('9.1 Progressive Discipline Steps')
    add_body('The standard progressive discipline process consists of four steps. However, the company reserves the right to skip steps or proceed directly to termination in cases of serious misconduct.')

    add_list_item('Verbal Warning: Informal discussion documented with a brief written summary')
    add_list_item('Written Warning: Formal documentation specifying the issue, expected improvement, and timeline')
    add_list_item('Final Written Warning: Last opportunity to correct behavior, typically with a 30-day improvement plan')
    add_list_item('Termination: Separation from the company when previous steps have not achieved improvement')
    add_blank_line()

    add_heading2('9.2 Grounds for Immediate Termination')
    add_body('Certain actions constitute gross misconduct and may result in immediate termination without prior progressive discipline. These include but are not limited to:')

    add_list_item('Theft, fraud, or embezzlement of company assets or funds')
    add_list_item('Physical violence, threats, or intimidation toward any person on company premises')
    add_list_item('Possession, distribution, or use of illegal substances on company property')
    add_list_item('Intentional destruction of company property or sabotage of operations')
    add_list_item('Unauthorized disclosure of Restricted or Confidential information')
    add_list_item('Falsification of employment records, time sheets, or expense reports')

    add_heading2('9.3 Appeals Process')
    add_body('Employees have the right to appeal any disciplinary action through the formal grievance procedure. Appeals must be submitted in writing to the Human Resources department within 10 business days of receiving the disciplinary notice.')
    add_body('The appeals process involves an independent review by a panel consisting of one HR representative, one management representative from outside the employee\'s department, and one peer representative selected by the employee. The panel\'s decision is final and binding.')

    doc.add_page_break()

    # ---- Section 10: Compliance and Ethics ----
    add_heading1('10. Compliance and Ethics')
    add_body('Compliance with all applicable laws, regulations, and internal policies is a fundamental requirement of employment at Meridian Global Solutions. The Chief Compliance Officer oversees the compliance program and reports directly to the Board of Directors.')

    add_heading2('10.1 Anti-Corruption and Bribery')
    add_body('The company strictly prohibits bribery, corruption, and improper payments in any form. This policy applies to all employees, agents, consultants, and third parties acting on behalf of the company, regardless of geographic location.')
    add_blank_line()
    add_body('Gifts and entertainment offered to or received from business contacts must comply with the Gift and Entertainment Policy. Gifts exceeding $100 in value must be reported to the Compliance team. Cash gifts are never acceptable regardless of amount.')

    add_heading2('10.2 Regulatory Compliance')
    add_body('Employees must comply with all industry-specific regulations applicable to their role and department. Mandatory compliance training is assigned based on job function and must be completed within the specified deadlines.')

    add_list_item('GDPR and data privacy requirements for employees handling EU personal data')
    add_list_item('SOX compliance obligations for finance and accounting personnel')
    add_list_item('HIPAA requirements for employees with access to protected health information')
    add_list_item('Export control regulations for teams involved in international product distribution')
    add_list_item('Anti-money laundering procedures for financial operations staff')

    add_heading2('10.3 Whistleblower Protection')
    add_body('The company maintains a robust whistleblower protection program. Employees who report suspected violations of law or company policy in good faith are protected from retaliation. Reports can be made through multiple channels including the ethics hotline, online portal, or directly to the Chief Compliance Officer.')
    add_blank_line()
    add_body('The Compliance team investigates all reported concerns and provides updates to reporters where legally permissible. The company cooperates fully with regulatory authorities in all investigations and audits.')

    doc.add_page_break()

    # ---- Appendix ----
    add_heading1('Appendix A: Key Contacts and Resources')

    add_heading2('Department Directory')
    add_body('Human Resources: ext. 4500, hr@meridianglobal.com. Information Technology: ext. 3100, it-help@meridianglobal.com. Security Operations: ext. 3300, security@meridianglobal.com. Legal Department: ext. 4200, legal@meridianglobal.com. Facilities Management: ext. 3200, facilities@meridianglobal.com.')
    add_blank_line()

    add_heading2('Emergency Contacts')
    add_body('Building Security (24/7): ext. 3333 or 1-888-555-GUARD. Ethics Hotline (anonymous): 1-888-555-ETHIC. IT Security Incident: 1-888-555-SECURE. Employee Assistance Program: 1-888-555-HELP.')
    add_body('For life-threatening emergencies, always call 911 first, then notify Building Security.')

    add_heading2('Policy Document Revision History')
    add_body('Version 4.2 (January 15, 2025): Updated remote work policy, added AI usage guidelines. Version 4.1 (July 1, 2024): Revised compensation structure, updated benefits enrollment periods. Version 4.0 (January 1, 2024): Major revision incorporating new compliance requirements. Version 3.5 (June 15, 2023): Added DEI section, updated travel expense limits.')
    add_blank_line()

    add_heading1('Appendix B: Acknowledgment Form')
    add_body('I acknowledge that I have received and read the Meridian Global Solutions Corporate Policy Manual (Version 4.2). I understand that it is my responsibility to comply with all policies and procedures described herein.')
    add_body('I understand that this manual does not constitute an employment contract and that the company reserves the right to modify these policies at any time with appropriate notice to employees.')
    add_blank_line()
    add_body('Employee Name: _______________________________________')
    add_body('Employee ID: ________________________________________')
    add_body('Department: _________________________________________')
    add_body('Date: _______________________________________________')
    add_body('Signature: __________________________________________')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
