"""
Reward Script: Apply consistent paragraph spacing to Corporate_Policy.docx
Task ID: writer_pd_043
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Heading 1 spacing — 24pt before, 12pt after
  Component 2 (0.25): Heading 2 spacing — 18pt before, 8pt after
  Component 3 (0.20): Body text spacing — 0pt before, 6pt after
  Component 4 (0.15): Body text alignment (justified) + 0.5cm first-line indent
  Component 5 (0.15): List items spacing — 0pt before, 3pt after
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_043'

# EMU constants for expected values
EMU_24PT = Pt(24)  # 304800
EMU_12PT = Pt(12)  # 152400
EMU_18PT = Pt(18)  # 228600
EMU_8PT = Pt(8)    # 101600
EMU_6PT = Pt(6)    # 76200
EMU_3PT = Pt(3)    # 38100
EMU_0PT = 0
EMU_05CM = Cm(0.5)  # 180000

# Tolerance for spacing comparison (allow ~1pt tolerance = 12700 EMU)
TOLERANCE = 15000


def emu_val(v):
    """Convert a spacing/indent value to integer EMU, treating None as 0 for inherited."""
    if v is None:
        return None  # Keep None to distinguish inherited vs explicit
    return int(v)


def spacing_match(actual, expected_emu):
    """Check if actual spacing (EMU) matches expected within tolerance."""
    if actual is None:
        # None means inherited from style; for paragraph-level overrides we need explicit
        return expected_emu == 0  # None could mean 0 (no override) only if expected is 0
    return abs(int(actual) - expected_emu) <= TOLERANCE


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Categorize paragraphs by style
    heading1_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 1']
    heading2_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 2']
    # Body text = Normal paragraphs with actual text, excluding the title line and subtitle
    normal_paras = [p for p in doc.paragraphs
                    if p.style and p.style.name == 'Normal'
                    and p.text.strip()
                    and p.text.strip() != 'Corporate Policy Manual'
                    and 'Version 4.2' not in p.text]
    list_paras = [p for p in doc.paragraphs
                  if p.style and p.style.name in ('List Bullet', 'List Number', 'List Paragraph')]

    print(f"Found: {len(heading1_paras)} H1, {len(heading2_paras)} H2, {len(normal_paras)} body, {len(list_paras)} list")

    # Component 1: Heading 1 spacing — 24pt before, 12pt after (0.25 points)
    try:
        if len(heading1_paras) == 0:
            print("FAIL: Component 1 — No Heading 1 paragraphs found")
        else:
            h1_correct = 0
            for p in heading1_paras:
                pf = p.paragraph_format
                sb_ok = spacing_match(pf.space_before, int(EMU_24PT))
                sa_ok = spacing_match(pf.space_after, int(EMU_12PT))
                if sb_ok and sa_ok:
                    h1_correct += 1
                else:
                    print(f"  H1 mismatch: sb={pf.space_before} (exp {int(EMU_24PT)}), sa={pf.space_after} (exp {int(EMU_12PT)}), text={p.text[:40]!r}")

            ratio = h1_correct / len(heading1_paras)
            if ratio >= 0.9:
                print(f"PASS: Component 1 — Heading 1 spacing correct ({h1_correct}/{len(heading1_paras)}) (0.25 pts)")
                total_score += 0.25
            elif ratio >= 0.5:
                pts = 0.25 * ratio
                print(f"PARTIAL: Component 1 — Heading 1 spacing {h1_correct}/{len(heading1_paras)} correct ({pts:.2f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 1 — Heading 1 spacing only {h1_correct}/{len(heading1_paras)} correct")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Heading 2 spacing — 18pt before, 8pt after (0.25 points)
    try:
        if len(heading2_paras) == 0:
            print("FAIL: Component 2 — No Heading 2 paragraphs found")
        else:
            h2_correct = 0
            for p in heading2_paras:
                pf = p.paragraph_format
                sb_ok = spacing_match(pf.space_before, int(EMU_18PT))
                sa_ok = spacing_match(pf.space_after, int(EMU_8PT))
                if sb_ok and sa_ok:
                    h2_correct += 1
                else:
                    print(f"  H2 mismatch: sb={pf.space_before} (exp {int(EMU_18PT)}), sa={pf.space_after} (exp {int(EMU_8PT)}), text={p.text[:40]!r}")

            ratio = h2_correct / len(heading2_paras)
            if ratio >= 0.9:
                print(f"PASS: Component 2 — Heading 2 spacing correct ({h2_correct}/{len(heading2_paras)}) (0.25 pts)")
                total_score += 0.25
            elif ratio >= 0.5:
                pts = 0.25 * ratio
                print(f"PARTIAL: Component 2 — Heading 2 spacing {h2_correct}/{len(heading2_paras)} correct ({pts:.2f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 2 — Heading 2 spacing only {h2_correct}/{len(heading2_paras)} correct")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body text spacing — 0pt before, 6pt after (0.20 points)
    try:
        if len(normal_paras) == 0:
            print("FAIL: Component 3 — No body text paragraphs found")
        else:
            body_correct = 0
            for p in normal_paras:
                pf = p.paragraph_format
                sb = pf.space_before
                sa = pf.space_after
                # space_before should be 0 (could be None=inherited or explicit 0)
                sb_ok = (sb is None or int(sb) <= TOLERANCE)
                sa_ok = spacing_match(sa, int(EMU_6PT))
                if sb_ok and sa_ok:
                    body_correct += 1

            ratio = body_correct / len(normal_paras)
            if ratio >= 0.9:
                print(f"PASS: Component 3 — Body text spacing correct ({body_correct}/{len(normal_paras)}) (0.20 pts)")
                total_score += 0.20
            elif ratio >= 0.5:
                pts = 0.20 * ratio
                print(f"PARTIAL: Component 3 — Body text spacing {body_correct}/{len(normal_paras)} correct ({pts:.2f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 3 — Body text spacing only {body_correct}/{len(normal_paras)} correct")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Body text justified alignment + 0.5cm first-line indent (0.15 points)
    try:
        if len(normal_paras) == 0:
            print("FAIL: Component 4 — No body text paragraphs found")
        else:
            align_correct = 0
            indent_correct = 0
            for p in normal_paras:
                pf = p.paragraph_format
                if pf.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    align_correct += 1
                fli = pf.first_line_indent
                if fli is not None and abs(int(fli) - int(EMU_05CM)) <= TOLERANCE:
                    indent_correct += 1

            align_ratio = align_correct / len(normal_paras)
            indent_ratio = indent_correct / len(normal_paras)

            sub_score = 0.0
            if align_ratio >= 0.9:
                sub_score += 0.075
                print(f"  PASS: Justified alignment ({align_correct}/{len(normal_paras)})")
            else:
                print(f"  FAIL: Justified alignment only {align_correct}/{len(normal_paras)}")

            if indent_ratio >= 0.9:
                sub_score += 0.075
                print(f"  PASS: First-line indent ({indent_correct}/{len(normal_paras)})")
            else:
                print(f"  FAIL: First-line indent only {indent_correct}/{len(normal_paras)}")

            if sub_score > 0:
                print(f"PASS: Component 4 — Alignment & indent ({sub_score:.3f} pts)")
                total_score += sub_score
            else:
                print(f"FAIL: Component 4 — Alignment & indent")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: List items spacing — 0pt before, 3pt after (0.15 points)
    try:
        if len(list_paras) == 0:
            print("FAIL: Component 5 — No list paragraphs found")
        else:
            list_correct = 0
            for p in list_paras:
                pf = p.paragraph_format
                sb = pf.space_before
                sa = pf.space_after
                sb_ok = (sb is None or int(sb) <= TOLERANCE)
                sa_ok = spacing_match(sa, int(EMU_3PT))
                if sb_ok and sa_ok:
                    list_correct += 1

            ratio = list_correct / len(list_paras)
            if ratio >= 0.9:
                print(f"PASS: Component 5 — List item spacing correct ({list_correct}/{len(list_paras)}) (0.15 pts)")
                total_score += 0.15
            elif ratio >= 0.5:
                pts = 0.15 * ratio
                print(f"PARTIAL: Component 5 — List item spacing {list_correct}/{len(list_paras)} correct ({pts:.2f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 5 — List item spacing only {list_correct}/{len(list_paras)} correct")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook — save any unsaved LibreOffice changes
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Main execution
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
