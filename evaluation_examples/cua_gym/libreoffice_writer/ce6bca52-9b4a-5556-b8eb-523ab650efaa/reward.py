"""
FINAL REWARD SCRIPT - SUCCESS
Task: LibreOffice Writer keeps littering my report with plain hyphens in ranges (e.g., 10-12, 1999-2003). I want every one of those to become a real en dash (–). What’s the quickest way to do a document-wide Find & Replace using the regular expression (?<=\d)-( ?)(?=\d) and swap it for the single en dash character, while leaving word hyphenation like "well-known" untouched?
Generated: 2025-09-10 17:21:41
Status: success
Model: azure-o3
Total Steps: 2
"""

import os
import re
from docx import Document


def _extract_all_text(doc: Document) -> str:
    """Collect all text from paragraphs and table cells into one string."""
    texts = []

    # Paragraphs outside tables
    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)

    # Paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        texts.append(p.text)

    return "\n".join(texts)


def verify_en_dash_replacement(file_path: str) -> float:
    """Verify that hyphens in numeric ranges were replaced with en dashes and
    that regular word-hyphenation remains unchanged.

    Scoring (progressive):
        • 0.4  – No numeric range still contains a plain hyphen (e.g. 10-12)
        • 0.3  – At least one numeric range now uses an en dash (e.g. 10–12)
        • 0.3  – Word hyphenations (e.g. well-known) still use a hyphen and did
                  not change to en dash
    Returns a float between 0.0 and 1.0 (exactly 1.0 for perfect completion).
    """

    score = 0.0
    max_score = 1.0

    # ----- 1. Load document -----
    if not os.path.exists(file_path):
        print("✗ File not found:", file_path)
        return 0.0  # cannot proceed

    try:
        doc = Document(file_path)
    except Exception as e:
        print("✗ Unable to open document:", e)
        return 0.0

    # ----- 2. Collect all text -----
    full_text = _extract_all_text(doc)

    # ----- 3. Identify patterns -----
    hyphen_numeric_ranges = re.findall(r"\d-\s?\d", full_text)
    en_dash_numeric_ranges = re.findall(r"\d–\s?\d", full_text)

    word_hyphenations_with_hyphen = re.findall(r"[A-Za-z]+-[A-Za-z]+", full_text)
    word_tokens_with_en_dash = re.findall(r"[A-Za-z]+–[A-Za-z]+", full_text)

    # Debug prints ---------------------------------------------------------
    print(f"Numeric ranges with hyphen       : {len(hyphen_numeric_ranges)} -> {hyphen_numeric_ranges[:10]}")
    print(f"Numeric ranges with en dash      : {len(en_dash_numeric_ranges)} -> {en_dash_numeric_ranges[:10]}")
    print(f"Word hyphenations using hyphen   : {len(word_hyphenations_with_hyphen)} -> {word_hyphenations_with_hyphen[:10]}")
    print(f"Word tokens using en dash        : {len(word_tokens_with_en_dash)} -> {word_tokens_with_en_dash[:10]}")
    print("--------------------------------------------------------------------")

    # ----- 4. Scoring criteria -----
    # 4a. All numeric ranges should be free of plain hyphens
    if len(hyphen_numeric_ranges) == 0:
        print("✓ No numeric ranges with plain hyphen found (0.4 points)")
        score += 0.4
    else:
        print("✗ Some numeric ranges still use plain hyphen (0 points)")

    # 4b. At least one numeric range now uses an en dash
    if len(en_dash_numeric_ranges) > 0:
        print("✓ Numeric ranges with en dash detected (0.3 points)")
        score += 0.3
    else:
        print("✗ No numeric ranges with en dash found (0 points)")

    # 4c. Word hyphenations should remain hyphens (not en dashes)
    if len(word_hyphenations_with_hyphen) > 0 and len(word_tokens_with_en_dash) == 0:
        print("✓ Word hyphenations correctly preserved with hyphen (0.3 points)")
        score += 0.3
    elif len(word_hyphenations_with_hyphen) > 0:
        # Some hyphenations preserved, but some were wrongly changed
        print("ⓘ Word hyphenations partly preserved, but some changed to en dash (0.15 points)")
        score += 0.15
    else:
        print("✗ No word hyphenations with hyphen detected (0 points)")

    # ----- 5. Final score -----
    final_score = min(score, max_score)
    print(f"Total Score: {final_score}")
    return final_score


if __name__ == "__main__":
    # Path to the document to be verified (provided by the environment)
    DOC_PATH = "/home/user/libreoffice_writer_keeps_littering_my_report_with_plain_hyphens_in_ranges_eg_10_12_1999_2003_i_want_.docx"

    reward_value = verify_en_dash_replacement(DOC_PATH)
    print(f"REWARD: {reward_value}")
