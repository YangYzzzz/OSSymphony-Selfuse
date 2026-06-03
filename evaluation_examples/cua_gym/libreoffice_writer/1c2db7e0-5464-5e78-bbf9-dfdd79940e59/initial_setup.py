"""
Initial Setup: Create a Writer document with tracked changes and comments for review extraction.
Task ID: writer_lec_095
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_095'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Namespace map for Word XML
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Three reviewers
REVIEWERS = [
    {"name": "Elena Rodriguez", "initials": "ER"},
    {"name": "James Mitchell", "initials": "JM"},
    {"name": "Priya Sharma", "initials": "PS"},
]

# Tracked change dates (spread across review period)
CHANGE_DATES = [
    "2025-11-03T09:15:00Z",
    "2025-11-03T09:22:00Z",
    "2025-11-03T10:05:00Z",
    "2025-11-04T11:30:00Z",
    "2025-11-04T14:10:00Z",
    "2025-11-04T14:45:00Z",
    "2025-11-05T08:50:00Z",
    "2025-11-05T09:30:00Z",
    "2025-11-05T10:15:00Z",
    "2025-11-06T13:20:00Z",
    "2025-11-06T14:00:00Z",
    "2025-11-06T16:30:00Z",
    "2025-11-07T09:00:00Z",
    "2025-11-07T10:45:00Z",
    "2025-11-07T11:30:00Z",
]

COMMENT_DATES = [
    "2025-11-03T09:18:00Z",
    "2025-11-03T10:10:00Z",
    "2025-11-04T11:35:00Z",
    "2025-11-04T14:50:00Z",
    "2025-11-05T09:00:00Z",
    "2025-11-05T10:20:00Z",
    "2025-11-06T14:05:00Z",
    "2025-11-07T09:10:00Z",
]


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, rpr_elem=None):
    """Create a w:r element with optional run properties."""
    r = OxmlElement('w:r')
    if rpr_elem is not None:
        r.append(copy.deepcopy(rpr_elem))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def add_insertion(para, text, author, date, change_id):
    """Add an insertion tracked change to a paragraph."""
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(change_id))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)
    r = make_run_element(text)
    ins.append(r)
    para._element.append(ins)


def add_deletion(para, text, author, date, change_id):
    """Add a deletion tracked change to a paragraph."""
    dele = OxmlElement('w:del')
    dele.set(qn('w:id'), str(change_id))
    dele.set(qn('w:author'), author)
    dele.set(qn('w:date'), date)
    r = OxmlElement('w:r')
    dt = OxmlElement('w:delText')
    dt.set(qn('xml:space'), 'preserve')
    dt.text = text
    r.append(dt)
    dele.append(r)
    para._element.append(dele)


def add_comment_to_paragraph(doc, para, comment_text, author, initials, date, comment_id, ref_text_start_idx=0):
    """Add a comment referencing text in a paragraph."""
    # Add comment to comments part
    comments_part = None
    for rel in doc.part.rels.values():
        if 'comments' in rel.reltype:
            comments_part = rel.target_part
            break

    # Add commentRangeStart
    range_start = OxmlElement('w:commentRangeStart')
    range_start.set(qn('w:id'), str(comment_id))

    # Add commentRangeEnd
    range_end = OxmlElement('w:commentRangeEnd')
    range_end.set(qn('w:id'), str(comment_id))

    # Add commentReference in a run
    comment_ref_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rstyle = OxmlElement('w:rStyle')
    rstyle.set(qn('w:val'), 'CommentReference')
    rpr.append(rstyle)
    comment_ref_run.append(rpr)
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), str(comment_id))
    comment_ref_run.append(ref)

    # Insert markers around text in paragraph
    p_elem = para._element
    runs = p_elem.findall(qn('w:r'))
    if runs:
        # Insert range start before first run, range end after last run
        p_elem.insert(list(p_elem).index(runs[0]), range_start)
        p_elem.append(range_end)
        p_elem.append(comment_ref_run)
    else:
        p_elem.append(range_start)
        p_elem.append(range_end)
        p_elem.append(comment_ref_run)

    return comment_id


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ===== DOCUMENT CONTENT: Project Proposal for CloudSync Platform =====

    # Title
    title = doc.add_heading('CloudSync Platform — Project Proposal', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)

    p1 = doc.add_paragraph()
    p1.add_run('CloudSync Platform is a next-generation cloud infrastructure management solution '
               'designed to unify multi-cloud deployments across AWS, Azure, and Google Cloud Platform. '
               'The platform provides centralized monitoring, automated scaling, and compliance reporting '
               'for enterprise customers managing distributed workloads.')

    p2 = doc.add_paragraph()
    p2.add_run('Our target market includes mid-size to large enterprises currently spending over $500,000 '
               'annually on cloud infrastructure. ')
    # Tracked Change 1: Elena inserted "with projected annual savings of 15-25%"
    add_insertion(p2, 'with projected annual savings of 15-25% ', REVIEWERS[0]["name"],
                  CHANGE_DATES[0], 101)
    p2.add_run('The initial deployment phase targets 50 pilot customers across financial services '
               'and healthcare sectors.')

    # Tracked Change 2: James deleted "across financial services and healthcare sectors"
    p3 = doc.add_paragraph()
    p3.add_run('The development timeline spans 18 months with three major release milestones. ')
    add_deletion(p3, 'We anticipate full market readiness by Q3 2026. ',
                 REVIEWERS[1]["name"], CHANGE_DATES[1], 102)
    p3.add_run('The revised schedule targets Q4 2026 for general availability.')

    # Section 2: Technical Architecture
    doc.add_heading('2. Technical Architecture', level=1)

    p4 = doc.add_paragraph()
    p4.add_run('The CloudSync architecture is built on a microservices foundation using Kubernetes '
               'orchestration. Each service module handles a specific domain: resource provisioning, '
               'monitoring, cost optimization, and security compliance.')

    # Tracked Change 3: Priya inserted "with automatic failover and disaster recovery"
    p5 = doc.add_paragraph()
    p5.add_run('The data layer utilizes a distributed PostgreSQL cluster ')
    add_insertion(p5, 'with automatic failover and disaster recovery ',
                  REVIEWERS[2]["name"], CHANGE_DATES[2], 103)
    p5.add_run('spanning three availability zones per region.')

    # Tracked Change 4: Elena deleted old API description
    p6 = doc.add_paragraph()
    p6.add_run('API gateway services handle authentication, rate limiting, and request routing. ')
    add_deletion(p6, 'All APIs follow REST conventions with JSON payloads. ',
                 REVIEWERS[0]["name"], CHANGE_DATES[3], 104)
    p6.add_run('The platform supports both REST and GraphQL interfaces with protocol buffer serialization.')

    # Tracked Change 5: James inserted performance requirement
    p7 = doc.add_paragraph()
    p7.add_run('Performance benchmarks indicate the system can handle ')
    add_insertion(p7, 'up to 10,000 concurrent API requests with sub-100ms latency ',
                  REVIEWERS[1]["name"], CHANGE_DATES[4], 105)
    p7.add_run('under standard load conditions.')

    # Section 3: Implementation Plan
    doc.add_heading('3. Implementation Plan', level=1)

    p8 = doc.add_paragraph()
    p8.add_run('Phase 1 (Months 1-6): Core infrastructure setup including Kubernetes cluster deployment, '
               'CI/CD pipeline configuration, and base monitoring framework.')

    # Tracked Change 6: Priya deleted outdated timeline detail
    p9 = doc.add_paragraph()
    p9.add_run('Phase 2 (Months 7-12): ')
    add_deletion(p9, 'Integration with third-party cloud providers and basic dashboard. ',
                 REVIEWERS[2]["name"], CHANGE_DATES[5], 106)
    p9.add_run('Advanced multi-cloud integration with unified dashboard and alerting system.')

    # Tracked Change 7: Elena inserted compliance detail
    p10 = doc.add_paragraph()
    p10.add_run('Phase 3 (Months 13-18): Enterprise features including ')
    add_insertion(p10, 'SOC 2 Type II and HIPAA compliance modules, ',
                  REVIEWERS[0]["name"], CHANGE_DATES[6], 107)
    p10.add_run('advanced analytics, and customer onboarding automation.')

    # Section 4: Budget and Resources
    doc.add_heading('4. Budget and Resources', level=1)

    p11 = doc.add_paragraph()
    p11.add_run('Total project budget: $2.4 million over 18 months. This includes personnel costs, '
                'infrastructure expenses, and third-party service subscriptions.')

    # Tracked Change 8: James deleted old budget line
    p12 = doc.add_paragraph()
    add_deletion(p12, 'Development team: 8 engineers at an average cost of $150,000 annually. ',
                 REVIEWERS[1]["name"], CHANGE_DATES[7], 108)
    p12.add_run('Development team: 12 engineers including 4 senior architects, 6 full-stack developers, '
                'and 2 DevOps specialists.')

    # Tracked Change 9: Priya inserted infrastructure cost detail
    p13 = doc.add_paragraph()
    p13.add_run('Infrastructure costs are estimated at $45,000 per month ')
    add_insertion(p13, 'including reserved instance pricing and enterprise support agreements ',
                  REVIEWERS[2]["name"], CHANGE_DATES[8], 109)
    p13.add_run('for the production environment.')

    # Section 5: Risk Assessment
    doc.add_heading('5. Risk Assessment', level=1)

    # Tracked Change 10: Elena inserted new risk item
    p14 = doc.add_paragraph()
    p14.add_run('Key risks identified during the planning phase include: vendor lock-in, talent retention, '
                'and regulatory changes. ')
    add_insertion(p14, 'Additionally, supply chain disruptions may impact hardware procurement timelines. ',
                  REVIEWERS[0]["name"], CHANGE_DATES[9], 110)

    # Tracked Change 11: James deleted old mitigation
    p15 = doc.add_paragraph()
    p15.add_run('Mitigation strategies: ')
    add_deletion(p15, 'Regular vendor reviews and contract renegotiation. ',
                 REVIEWERS[1]["name"], CHANGE_DATES[10], 111)
    p15.add_run('Multi-vendor qualification process with automatic failover between cloud providers.')

    # Tracked Change 12: Priya inserted security consideration
    p16 = doc.add_paragraph()
    p16.add_run('Security considerations require ')
    add_insertion(p16, 'quarterly penetration testing and annual third-party security audits ',
                  REVIEWERS[2]["name"], CHANGE_DATES[11], 112)
    p16.add_run('as part of the ongoing operational budget.')

    # Section 6: Success Metrics
    doc.add_heading('6. Success Metrics', level=1)

    p17 = doc.add_paragraph()
    p17.add_run('Platform adoption targets: 50 customers in Year 1, scaling to 200 by end of Year 2.')

    # Tracked Change 13: Elena deleted old metric
    p18 = doc.add_paragraph()
    add_deletion(p18, 'Customer satisfaction target: 85% positive NPS score. ',
                 REVIEWERS[0]["name"], CHANGE_DATES[12], 113)
    p18.add_run('Customer satisfaction target: 90% positive NPS score with quarterly survey cadence.')

    # Tracked Change 14: James inserted revenue target
    p19 = doc.add_paragraph()
    p19.add_run('Revenue projections: ')
    add_insertion(p19, '$1.2 million ARR by end of Year 1, growing to $4.8 million ARR by Year 3 ',
                  REVIEWERS[1]["name"], CHANGE_DATES[13], 114)
    p19.add_run('based on current pipeline analysis.')

    # Tracked Change 15: Priya deleted outdated KPI
    p20 = doc.add_paragraph()
    p20.add_run('Technical KPIs include 99.95% uptime SLA ')
    add_deletion(p20, 'and average response time under 500ms',
                 REVIEWERS[2]["name"], CHANGE_DATES[14], 115)
    p20.add_run('and average response time under 200ms with P99 latency below 500ms.')

    # Conclusion
    doc.add_heading('7. Conclusion', level=1)
    p_last = doc.add_paragraph()
    p_last.add_run('CloudSync Platform represents a significant market opportunity in the rapidly growing '
                   'cloud management space. With the right investment and execution, we are positioned to '
                   'capture meaningful market share and deliver substantial returns.')

    # ===== ADD COMMENTS via XML manipulation =====
    # We need to create a comments part and add comments to it
    _add_comments_part(doc)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def _add_comments_part(doc):
    """Add a comments XML part to the document with 8 comments from 3 reviewers."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    # Comment definitions: (comment_id, author, initials, date, comment_text)
    comments_data = [
        (201, REVIEWERS[0]["name"], REVIEWERS[0]["initials"], COMMENT_DATES[0],
         "The savings projection should be backed by case studies from comparable deployments."),
        (202, REVIEWERS[2]["name"], REVIEWERS[2]["initials"], COMMENT_DATES[1],
         "Consider adding a diagram showing the failover architecture between availability zones."),
        (203, REVIEWERS[0]["name"], REVIEWERS[0]["initials"], COMMENT_DATES[2],
         "The REST vs GraphQL decision should be documented in the architecture decision records."),
        (204, REVIEWERS[1]["name"], REVIEWERS[1]["initials"], COMMENT_DATES[3],
         "These performance numbers need to be validated with load testing before we commit to them publicly."),
        (205, REVIEWERS[2]["name"], REVIEWERS[2]["initials"], COMMENT_DATES[4],
         "Phase 2 timeline seems aggressive. Can we add a buffer for integration testing?"),
        (206, REVIEWERS[1]["name"], REVIEWERS[1]["initials"], COMMENT_DATES[5],
         "SOC 2 compliance certification typically takes 6-9 months. Factor this into the Phase 3 schedule."),
        (207, REVIEWERS[0]["name"], REVIEWERS[0]["initials"], COMMENT_DATES[6],
         "The security audit budget should be itemized separately from the infrastructure costs."),
        (208, REVIEWERS[1]["name"], REVIEWERS[1]["initials"], COMMENT_DATES[7],
         "NPS target of 90% is ambitious for a new platform. Consider setting interim milestones."),
    ]

    # Build comments XML
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">\n'
    )

    for cid, author, initials, date, text in comments_data:
        comments_xml += (
            f'  <w:comment w:id="{cid}" w:author="{author}" w:initials="{initials}" w:date="{date}">\n'
            f'    <w:p>\n'
            f'      <w:r>\n'
            f'        <w:t>{text}</w:t>\n'
            f'      </w:r>\n'
            f'    </w:p>\n'
            f'  </w:comment>\n'
        )
    comments_xml += '</w:comments>'

    # Create the comments part
    comments_part_name = '/word/comments.xml'
    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'

    part = Part(
        PackURI(comments_part_name),
        content_type,
        comments_xml.encode('utf-8'),
        doc.part.package,
    )

    # Add relationship from document part to comments part
    COMMENTS_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    doc.part.relate_to(part, COMMENTS_REL_TYPE)

    # Now add comment range markers in the document body for each comment
    # We'll attach comments to specific paragraphs
    paragraphs = doc.paragraphs
    # Map comment IDs to paragraph indices where they should appear
    comment_para_map = {
        201: 2,   # Executive Summary - savings projection paragraph
        202: 5,   # Technical Architecture - distributed DB paragraph
        203: 6,   # API gateway paragraph
        204: 7,   # Performance benchmarks paragraph
        205: 9,   # Phase 2 paragraph
        206: 10,  # Phase 3 paragraph
        207: 16,  # Security considerations paragraph
        208: 18,  # NPS score paragraph
    }

    for cid, para_idx in comment_para_map.items():
        if para_idx < len(paragraphs):
            para = paragraphs[para_idx]
            p_elem = para._element

            # Add commentRangeStart before first child
            range_start = OxmlElement('w:commentRangeStart')
            range_start.set(qn('w:id'), str(cid))
            p_elem.insert(0, range_start)

            # Add commentRangeEnd and reference at end
            range_end = OxmlElement('w:commentRangeEnd')
            range_end.set(qn('w:id'), str(cid))
            p_elem.append(range_end)

            # Comment reference run
            ref_run = OxmlElement('w:r')
            ref_rpr = OxmlElement('w:rPr')
            ref_style = OxmlElement('w:rStyle')
            ref_style.set(qn('w:val'), 'CommentReference')
            ref_rpr.append(ref_style)
            ref_run.append(ref_rpr)
            ref_elem = OxmlElement('w:commentReference')
            ref_elem.set(qn('w:id'), str(cid))
            ref_run.append(ref_elem)
            p_elem.append(ref_run)


create_initial()
