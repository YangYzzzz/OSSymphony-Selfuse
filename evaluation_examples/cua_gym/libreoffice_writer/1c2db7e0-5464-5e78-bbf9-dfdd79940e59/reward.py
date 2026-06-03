"""
Reward Script: Export comments and tracked changes to a review summary document
Task ID: writer_lec_095
Domain: libreoffice_writer
Scoring:
  C1 (0.15) — review_summary.docx exists on Desktop and is loadable
  C2 (0.15) — Contains "Tracked Changes" and "Comments" section headings
  C3 (0.25) — Contains 15 tracked-change tables with 5-row structure (Author/Date/Type/Original/Modified)
  C4 (0.20) — Contains 8 comment tables with 4-row structure (Author/Date/Referenced Text/Comment)
  C5 (0.15) — All 3 reviewers appear across entries (Elena Rodriguez, James Mitchell, Priya Sharma)
  C6 (0.10) — Change types include both Insertion and Deletion entries
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_095'
SUMMARY_PATH = os.path.join(WORKDIR, 'Desktop', 'review_summary.docx')


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Component 1: review_summary.docx exists on Desktop and is loadable (0.15 pts)
    # This is a task-introduced file — it should NOT exist in initial_env.
    try:
        if not os.path.exists(SUMMARY_PATH):
            print(f"FAIL: Component 1 — File not found at {SUMMARY_PATH}")
            print("REWARD: 0.0")
            return 0.0

        from docx import Document
        doc = Document(SUMMARY_PATH)
        if len(doc.paragraphs) > 0:
            print(f"PASS: Component 1 — review_summary.docx loaded ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — File loaded but is empty (0 paragraphs)")
    except Exception as e:
        print(f"ERROR: Component 1 — Cannot load file: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 2: Contains "Tracked Changes" and "Comments" section headings (0.15 pts)
    try:
        para_texts_lower = [p.text.strip().lower() for p in doc.paragraphs]
        has_tracked = any('tracked changes' in t for t in para_texts_lower)
        has_comments = any('comments' in t for t in para_texts_lower)

        if has_tracked and has_comments:
            print(f"PASS: Component 2 — Both 'Tracked Changes' and 'Comments' headings found (0.15 pts)")
            total_score += 0.15
        elif has_tracked:
            print(f"PARTIAL: Component 2 — 'Tracked Changes' found but 'Comments' missing (0.075 pts)")
            total_score += 0.075
        elif has_comments:
            print(f"PARTIAL: Component 2 — 'Comments' found but 'Tracked Changes' missing (0.075 pts)")
            total_score += 0.075
        else:
            print(f"FAIL: Component 2 — Neither 'Tracked Changes' nor 'Comments' heading found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Classify tables into change-tables (5-row) and comment-tables (4-row)
    # Change tables have rows: Author, Date, Type, Original Text, Modified Text
    # Comment tables have rows: Author, Date, Referenced Text, Comment
    change_tables = []
    comment_tables = []
    all_authors = set()
    change_types = set()

    try:
        for table in doc.tables:
            rows = table.rows
            num_rows = len(rows)
            # Extract first-column labels
            labels = []
            for row in rows:
                cell_text = row.cells[0].text.strip().lower()
                labels.append(cell_text)

            if num_rows == 5:
                # Check if it looks like a change table
                has_author = any('author' in l for l in labels)
                has_date = any('date' in l for l in labels)
                has_type = any('type' in l for l in labels)
                has_orig = any('original' in l for l in labels)
                has_mod = any('modified' in l for l in labels)
                if has_author and has_date and has_type and (has_orig or has_mod):
                    change_tables.append(table)
                    # Extract author and type
                    for row in rows:
                        label = row.cells[0].text.strip().lower()
                        value = row.cells[1].text.strip()
                        if 'author' in label and value:
                            all_authors.add(value)
                        if 'type' in label and value:
                            change_types.add(value.lower())

            elif num_rows == 4:
                # Check if it looks like a comment table
                has_author = any('author' in l for l in labels)
                has_date = any('date' in l for l in labels)
                has_ref = any('reference' in l for l in labels)
                has_comment = any('comment' in l for l in labels)
                if has_author and has_date and (has_ref or has_comment):
                    comment_tables.append(table)
                    # Extract author
                    for row in rows:
                        label = row.cells[0].text.strip().lower()
                        value = row.cells[1].text.strip()
                        if 'author' in label and value:
                            all_authors.add(value)
    except Exception as e:
        print(f"ERROR: Table classification — {e}")

    # Component 3: 15 tracked-change tables (0.25 pts)
    try:
        num_changes = len(change_tables)
        if num_changes >= 15:
            print(f"PASS: Component 3 — Found {num_changes} tracked-change tables (expected 15) (0.25 pts)")
            total_score += 0.25
        elif num_changes >= 10:
            partial = round(0.25 * (num_changes / 15), 3)
            print(f"PARTIAL: Component 3 — Found {num_changes}/15 tracked-change tables ({partial} pts)")
            total_score += partial
        elif num_changes >= 5:
            partial = round(0.25 * (num_changes / 15), 3)
            print(f"PARTIAL: Component 3 — Found {num_changes}/15 tracked-change tables ({partial} pts)")
            total_score += partial
        elif num_changes > 0:
            partial = round(0.25 * (num_changes / 15), 3)
            print(f"PARTIAL: Component 3 — Found {num_changes}/15 tracked-change tables ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No tracked-change tables found (expected 15)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 8 comment tables (0.20 pts)
    try:
        num_comments = len(comment_tables)
        if num_comments >= 8:
            print(f"PASS: Component 4 — Found {num_comments} comment tables (expected 8) (0.20 pts)")
            total_score += 0.20
        elif num_comments >= 4:
            partial = round(0.20 * (num_comments / 8), 3)
            print(f"PARTIAL: Component 4 — Found {num_comments}/8 comment tables ({partial} pts)")
            total_score += partial
        elif num_comments > 0:
            partial = round(0.20 * (num_comments / 8), 3)
            print(f"PARTIAL: Component 4 — Found {num_comments}/8 comment tables ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No comment tables found (expected 8)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: All 3 reviewers appear (0.15 pts)
    try:
        expected_reviewers = {'Elena Rodriguez', 'James Mitchell', 'Priya Sharma'}
        found_reviewers = set()
        for author in all_authors:
            for expected in expected_reviewers:
                if expected.lower() in author.lower():
                    found_reviewers.add(expected)
        num_found = len(found_reviewers)
        if num_found == 3:
            print(f"PASS: Component 5 — All 3 reviewers found: {found_reviewers} (0.15 pts)")
            total_score += 0.15
        elif num_found > 0:
            partial = round(0.15 * (num_found / 3), 3)
            print(f"PARTIAL: Component 5 — Found {num_found}/3 reviewers: {found_reviewers} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — No expected reviewers found. Authors seen: {all_authors}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Change types include both Insertion and Deletion (0.10 pts)
    try:
        has_insertion = any('insert' in t for t in change_types)
        has_deletion = any('delet' in t for t in change_types)
        if has_insertion and has_deletion:
            print(f"PASS: Component 6 — Both Insertion and Deletion types found: {change_types} (0.10 pts)")
            total_score += 0.10
        elif has_insertion or has_deletion:
            print(f"PARTIAL: Component 6 — Only one change type found: {change_types} (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No Insertion/Deletion types found. Types seen: {change_types}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any open LibreOffice documents before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state()
verify_task()
