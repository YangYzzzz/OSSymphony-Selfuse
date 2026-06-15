"""
Reward Script: Add professional cover page to SaaS e-book
Task ID: writer_mktg_028
Domain: libreoffice_writer
Scoring:
  - Component 1: Cover title present, bold, ~28pt, centered (0.35 pts)
  - Component 2: Subtitle present, italic, ~16pt, centered (0.20 pts)
  - Component 3: Author and date present, centered (0.20 pts)
  - Component 4: Page break after cover + Chapter 1 is first body content (0.25 pts)
  Total: 1.0
"""

import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_028'
FILE_PATH = os.path.join(WORKDIR, 'saas_growth_ebook.docx')


def count_page_breaks(doc):
    """Count manual page breaks in the document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    count = 0
    break_indices = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                if br.attrib.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
                ) == 'page':
                    count += 1
                    break_indices.append(i)
    return count, break_indices


def find_paragraph_with_text(doc, target_text):
    """Return (index, paragraph) for the first paragraph containing target_text, or (None, None)."""
    for i, p in enumerate(doc.paragraphs):
        if target_text in p.text:
            return i, p
    return None, None


def is_centered(para):
    """Return True if paragraph alignment is CENTER."""
    return para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER


def verify_task(file_path):
    """
    Verify cover page was added to saas_growth_ebook.docx.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Cover title "The SaaS Growth Playbook" — bold, ~28pt, centered
    # (0.35 points)
    # Expected: paragraph with this text appears BEFORE Chapter 1, is centered,
    # has at least one run with bold=True and font size ~28pt.
    # -------------------------------------------------------------------------
    try:
        TITLE_TEXT = 'The SaaS Growth Playbook'
        title_idx, title_para = find_paragraph_with_text(doc, TITLE_TEXT)
        ch1_idx, ch1_para = find_paragraph_with_text(doc, 'Chapter 1:')

        if title_para is None:
            print("FAIL: Component 1 — Title '%s' not found in document" % TITLE_TEXT)
        elif ch1_idx is not None and title_idx >= ch1_idx:
            print("FAIL: Component 1 — Title found at para %d but Chapter 1 is at para %d (title must come before body)" % (title_idx, ch1_idx))
        else:
            # Check alignment
            centered = is_centered(title_para)
            # Check bold and font size on runs
            bold_found = any(r.font.bold for r in title_para.runs if r.text.strip())
            size_ok = any(
                r.font.size is not None and abs(r.font.size.pt - 28.0) < 2.0
                for r in title_para.runs if r.text.strip()
            )

            if centered and bold_found and size_ok:
                print("PASS: Component 1 — Title found at para %d, centered=%s, bold=%s, size~28pt" % (title_idx, centered, bold_found))
                total_score += 0.35
            elif centered and bold_found:
                # Partial: title is there and bold/centered but font size not exactly 28pt
                print("PARTIAL: Component 1 — Title found (centered, bold) but size check failed (size_ok=%s); awarding 0.20" % size_ok)
                total_score += 0.20
            else:
                print("FAIL: Component 1 — Title found at para %d but centered=%s, bold=%s, size_ok=%s" % (title_idx, centered, bold_found, size_ok))
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # -------------------------------------------------------------------------
    # Component 2: Subtitle "Strategies for Scaling from $1M to $10M ARR" —
    # italic, ~16pt, centered (0.20 points)
    # -------------------------------------------------------------------------
    try:
        SUBTITLE_TEXT = 'Strategies for Scaling from $1M to $10M ARR'
        sub_idx, sub_para = find_paragraph_with_text(doc, SUBTITLE_TEXT)
        ch1_idx, _ = find_paragraph_with_text(doc, 'Chapter 1:')

        if sub_para is None:
            print("FAIL: Component 2 — Subtitle '%s' not found" % SUBTITLE_TEXT)
        elif ch1_idx is not None and sub_idx >= ch1_idx:
            print("FAIL: Component 2 — Subtitle found at para %d but Chapter 1 is at para %d (must be before body)" % (sub_idx, ch1_idx))
        else:
            centered = is_centered(sub_para)
            italic_found = any(r.font.italic for r in sub_para.runs if r.text.strip())
            size_ok = any(
                r.font.size is not None and abs(r.font.size.pt - 16.0) < 2.0
                for r in sub_para.runs if r.text.strip()
            )

            if centered and italic_found and size_ok:
                print("PASS: Component 2 — Subtitle found at para %d, centered=%s, italic=%s, size~16pt" % (sub_idx, centered, italic_found))
                total_score += 0.20
            elif centered and italic_found:
                print("PARTIAL: Component 2 — Subtitle (centered, italic) but size not 16pt; awarding 0.10")
                total_score += 0.10
            else:
                print("FAIL: Component 2 — Subtitle found at para %d but centered=%s, italic=%s, size_ok=%s" % (sub_idx, centered, italic_found, size_ok))
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # -------------------------------------------------------------------------
    # Component 3: Author "Written by the Apex Dynamics Growth Team" and
    # date "March 2026" — both centered, present before Chapter 1 (0.20 points)
    # -------------------------------------------------------------------------
    try:
        AUTHOR_TEXT = 'Written by the Apex Dynamics Growth Team'
        DATE_TEXT = 'March 2026'
        ch1_idx, _ = find_paragraph_with_text(doc, 'Chapter 1:')

        author_idx, author_para = find_paragraph_with_text(doc, AUTHOR_TEXT)
        date_idx, date_para = find_paragraph_with_text(doc, DATE_TEXT)

        author_ok = (author_para is not None and
                     (ch1_idx is None or author_idx < ch1_idx) and
                     is_centered(author_para))
        date_ok = (date_para is not None and
                   (ch1_idx is None or date_idx < ch1_idx) and
                   is_centered(date_para))

        if author_ok and date_ok:
            print("PASS: Component 3 — Author at para %d and date at para %d, both centered before Chapter 1" % (author_idx, date_idx))
            total_score += 0.20
        elif author_ok:
            print("PARTIAL: Component 3 — Author found and centered, but date missing/wrong; awarding 0.10")
            total_score += 0.10
        elif date_ok:
            print("PARTIAL: Component 3 — Date found and centered, but author missing/wrong; awarding 0.10")
            total_score += 0.10
        else:
            print("FAIL: Component 3 — author_ok=%s (author_para=%s), date_ok=%s (date_para=%s)" % (
                author_ok, author_para is not None, date_ok, date_para is not None))
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    # -------------------------------------------------------------------------
    # Component 4: Page break after cover + Chapter 1 is first body paragraph
    # after the page break (0.25 points)
    # -------------------------------------------------------------------------
    try:
        num_breaks, break_indices = count_page_breaks(doc)
        ch1_idx, ch1_para = find_paragraph_with_text(doc, 'Chapter 1:')

        has_page_break = num_breaks >= 1
        # Page break paragraph must appear before Chapter 1
        break_before_ch1 = ch1_idx is not None and any(bi < ch1_idx for bi in break_indices)
        # Chapter 1 should be the first Heading 1 (or first chapter) after the break
        ch1_is_first_body = False
        if break_before_ch1:
            # Find the break para just before Chapter 1
            cover_break_idx = max(bi for bi in break_indices if bi < ch1_idx)
            # Check that Chapter 1 follows (within 2 paragraphs) after the break
            ch1_is_first_body = ch1_idx > cover_break_idx and ch1_idx <= cover_break_idx + 3

        if has_page_break and break_before_ch1 and ch1_is_first_body:
            print("PASS: Component 4 — Page break at para %d, Chapter 1 at para %d (first body content)" % (cover_break_idx, ch1_idx))
            total_score += 0.25
        elif has_page_break and break_before_ch1:
            print("PARTIAL: Component 4 — Page break before Chapter 1 but Chapter 1 not immediately after break; awarding 0.15")
            total_score += 0.15
        elif has_page_break:
            print("PARTIAL: Component 4 — Page break exists (at %s) but not before Chapter 1 (at para %s); awarding 0.05" % (break_indices, ch1_idx))
            total_score += 0.05
        else:
            print("FAIL: Component 4 — No page break found (breaks=%d)" % num_breaks)
    except Exception as e:
        print("ERROR: Component 4 — %s" % e)

    final_score = min(total_score, 1.0)
    print("\nScore: %s/1.0" % total_score)
    print("REWARD: %s" % final_score)
    return final_score


if not os.path.exists(FILE_PATH):
    print("File not found: %s" % FILE_PATH)
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
