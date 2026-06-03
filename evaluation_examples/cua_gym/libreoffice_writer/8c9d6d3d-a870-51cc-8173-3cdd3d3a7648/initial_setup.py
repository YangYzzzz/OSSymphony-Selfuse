"""
Initial Setup: SaaS Growth E-book without cover page
Task ID: writer_mktg_028
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'saas_growth_ebook'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Chapter 1: Product-Led Growth ---
    h1 = doc.add_heading('Chapter 1: Product-Led Growth', level=1)

    p = doc.add_paragraph(
        'Product-led growth (PLG) has emerged as one of the most powerful go-to-market '
        'strategies for SaaS companies targeting the $1M to $10M ARR milestone. Rather than '
        'relying solely on a direct sales force, PLG companies let the product itself drive '
        'user acquisition, retention, and expansion.'
    )

    p = doc.add_paragraph(
        'Companies like Slack, Dropbox, and Figma built their initial user bases primarily '
        'through organic product adoption. When users experience immediate value, they naturally '
        'invite colleagues and share the product within their organizations. This viral loop can '
        'dramatically reduce customer acquisition costs (CAC) while improving retention rates.'
    )

    p = doc.add_paragraph(
        'To implement PLG effectively, your product must deliver value before requiring payment. '
        'A generous free tier or trial period lets prospects experience core functionality. '
        'The key metrics to track are time-to-value (TTV), product qualified leads (PQLs), and '
        'activation rate — the percentage of users who complete the critical "aha moment" action.'
    )

    p = doc.add_paragraph(
        'Apex Dynamics implemented a freemium model in Q3 2024, offering unlimited projects for '
        'teams up to five users. Within six months, organic signups increased by 340%, and '
        'conversion from free to paid grew from 4.2% to 11.8% as onboarding was refined.'
    )

    # --- Chapter 2: Customer Success as a Growth Engine ---
    doc.add_heading('Chapter 2: Customer Success as a Growth Engine', level=1)

    p = doc.add_paragraph(
        'Many early-stage SaaS companies treat customer success (CS) as a cost center — a '
        'reactive support function. High-growth companies recognize CS as a revenue driver. '
        'A well-structured CS team reduces churn, identifies upsell opportunities, and turns '
        'customers into advocates who generate referrals.'
    )

    p = doc.add_paragraph(
        'The foundation of great customer success is proactive engagement. Rather than waiting '
        'for customers to file support tickets, CS teams should monitor product usage data and '
        'reach out when engagement drops or when customers hit friction points. Tools like '
        'Gainsight, ChurnZero, or even a simple Salesforce dashboard can surface at-risk accounts.'
    )

    p = doc.add_paragraph(
        'Net Revenue Retention (NRR) is the north star metric for customer success. An NRR above '
        '120% means your existing customers are growing faster than you lose revenue from churn. '
        'Best-in-class SaaS companies report NRR between 120% and 145%. Achieving this requires '
        'a deliberate expansion motion — tiered pricing, add-on modules, or seat expansion.'
    )

    p = doc.add_paragraph(
        'At Apex Dynamics, each CS manager owns a book of business of approximately 40 accounts. '
        'Monthly business reviews (MBRs) are mandatory for accounts above $10K ARR. The CS team '
        'contributed $620K in expansion revenue in fiscal year 2025, representing 28% of total '
        'new ARR for the year.'
    )

    # --- Chapter 3: Pricing Strategy and Packaging ---
    doc.add_heading('Chapter 3: Pricing Strategy and Packaging', level=1)

    p = doc.add_paragraph(
        'Pricing is one of the highest-leverage decisions in a SaaS company yet it receives '
        'surprisingly little analytical attention. Getting pricing right can accelerate revenue '
        'growth without adding a single new customer. Getting it wrong can stunt growth even '
        'when the product is excellent.'
    )

    p = doc.add_paragraph(
        'Value-based pricing — charging based on the value delivered rather than cost plus margin '
        '— is the gold standard for SaaS. To implement it, you must first quantify the economic '
        'impact of your product for representative customer segments. Surveys, win/loss analysis, '
        'and customer interviews are essential inputs.'
    )

    p = doc.add_paragraph(
        'Packaging decisions matter as much as price points. The classic Good-Better-Best structure '
        'works well for most SaaS products. Tier names like Starter, Professional, and Enterprise '
        'signal the target buyer. Feature gating should feel logical: put the features most valued '
        'by larger teams in higher tiers, not just arbitrarily restrict lower tiers.'
    )

    p = doc.add_paragraph(
        'Annual contracts provide predictable ARR and typically command a 15-20% discount over '
        'monthly billing. Pushing customers toward annual commitments early improves cash flow '
        'and reduces monthly churn risk. Offering a two-year prepay at a 25% discount can be '
        'attractive for budget-constrained buyers and provides a significant cash injection.'
    )

    # --- Chapter 4: Demand Generation at Scale ---
    doc.add_heading('Chapter 4: Demand Generation at Scale', level=1)

    p = doc.add_paragraph(
        'Scaling from $1M to $10M ARR requires building a repeatable demand generation engine. '
        'The channels and tactics that worked to reach the first million will not be sufficient '
        'to reach the next nine million. This chapter covers the demand generation strategies '
        'that growth-stage SaaS companies use most effectively.'
    )

    p = doc.add_paragraph(
        'Content marketing remains one of the most cost-effective channels for B2B SaaS. A '
        'well-executed SEO content strategy can drive 30-50% of inbound leads at a fraction of '
        'the cost of paid acquisition. The key is targeting high-intent search terms that '
        'indicate buying readiness, not just general informational queries.'
    )

    p = doc.add_paragraph(
        'Account-based marketing (ABM) is increasingly essential for SaaS companies with an '
        'average contract value (ACV) above $20K. ABM aligns sales and marketing around a '
        'defined list of target accounts and uses personalized campaigns to engage multiple '
        'stakeholders within each account simultaneously.'
    )

    p = doc.add_paragraph(
        'Partner channels — including technology integrations, resellers, and referral programs '
        '— can contribute 20-30% of new ARR for well-structured programs. Integration partners '
        'are particularly valuable: customers who use your product alongside complementary tools '
        'show 40% higher retention rates and 2.3x higher lifetime value.'
    )

    # --- Chapter 5: Sales Process Optimization ---
    doc.add_heading('Chapter 5: Sales Process Optimization', level=1)

    p = doc.add_paragraph(
        'As a SaaS company scales, informal sales processes become a liability. Deals are lost '
        'to poor follow-up, opportunities stall without clear next steps, and forecast accuracy '
        'suffers. Systematic sales process design is what separates companies that achieve '
        'consistent growth from those that plateau.'
    )

    p = doc.add_paragraph(
        'A well-defined sales methodology gives your team a common language and framework. '
        'MEDDIC (Metrics, Economic Buyer, Decision Criteria, Decision Process, Identify Pain, '
        'Champion) is widely used in enterprise SaaS. For mid-market and SMB, lighter frameworks '
        'like SPIN Selling or Challenger Sale can be more appropriate given shorter sales cycles.'
    )

    p = doc.add_paragraph(
        'Pipeline velocity is the core metric for sales operations: Pipeline Velocity = '
        '(Number of Opportunities × Win Rate × ACV) / Sales Cycle Length. Improving any of '
        'the four variables increases revenue. Sales ops teams should build dashboards that '
        'surface leading indicators for each variable and run structured experiments to improve them.'
    )

    p = doc.add_paragraph(
        'In 2025, Apex Dynamics rebuilt its sales qualification process using a combination of '
        'MEDDIC and product usage signals. Average sales cycle decreased from 47 days to 31 days. '
        'Win rate improved from 22% to 34%, and average deal size grew 18% as reps focused '
        'attention on accounts with the strongest economic justification.'
    )

    # --- Chapter 6: Building a Data-Driven Culture ---
    doc.add_heading('Chapter 6: Building a Data-Driven Culture', level=1)

    p = doc.add_paragraph(
        'The final chapter addresses what many consider the hardest challenge of scaling a SaaS '
        'business: building a truly data-driven culture. Data infrastructure is necessary but '
        'not sufficient. The bigger challenge is changing how teams make decisions — replacing '
        'intuition and HiPPO dynamics with structured experimentation and evidence-based reasoning.'
    )

    p = doc.add_paragraph(
        'Start with a metrics hierarchy that connects company-level OKRs to team-level KPIs and '
        'individual contributors\' daily activity metrics. When every team member can draw a line '
        'from their daily work to the company\'s revenue targets, focus and accountability improve '
        'dramatically. Tools like Looker, Tableau, or Metabase can democratize data access.'
    )

    p = doc.add_paragraph(
        'Experimentation culture requires psychological safety. Teams must feel empowered to test '
        'ideas, measure results honestly, and report failures without career risk. A simple A/B '
        'testing framework — hypothesis → test design → statistical significance → decision — '
        'should become standard operating procedure for product, marketing, and CS teams.'
    )

    p = doc.add_paragraph(
        'Regular business reviews at every level create accountability loops. Weekly team standups '
        'review leading indicators. Monthly business reviews examine lagging metrics and trend '
        'lines. Quarterly planning cycles use data to set targets and allocate resources. This '
        'cadence ensures decisions are grounded in evidence rather than anecdote.'
    )

    p = doc.add_paragraph(
        'The journey from $1M to $10M ARR is one of the most challenging and rewarding phases '
        'of building a SaaS company. The strategies in this playbook — product-led growth, '
        'customer success, value-based pricing, demand generation, sales excellence, and data '
        'culture — are not a checklist but an integrated system. Success requires executing all '
        'of them in concert, adapting continuously based on what your data tells you.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
