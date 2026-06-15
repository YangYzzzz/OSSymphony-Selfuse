"""
Initial Setup: Business report with three section headings in default Heading 1 style.
Task ID: writer_biz_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Quarterly Business Review — Q1 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This report provides a comprehensive overview of our business performance '
        'during the first quarter of 2025. The following sections cover market conditions, '
        'competitive positioning, and our forward-looking growth strategy.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # --- Section 1: Market Analysis ---
    doc.add_heading('Market Analysis', level=1)

    doc.add_paragraph(
        'The global enterprise software market reached $685 billion in Q1 2025, '
        'representing a 9.2% year-over-year increase. Cloud-native solutions continued '
        'to dominate new deployments, with 73% of enterprise buyers prioritizing SaaS '
        'platforms over on-premises alternatives.'
    )
    doc.add_paragraph(
        'Our primary addressable market in the Asia-Pacific region grew by 14.8%, '
        'outpacing both North America (7.1%) and Europe (5.9%). Key drivers included '
        'accelerated digital transformation initiatives in financial services, healthcare, '
        'and government sectors. The shift toward AI-augmented workflows created new '
        'demand segments worth an estimated $42 billion annually.'
    )
    doc.add_paragraph(
        'Customer acquisition costs decreased by 11% compared to Q4 2024, largely due '
        'to improvements in our inbound marketing funnel and expanded partner channel '
        'referrals. Average deal size increased to $128,500, up from $112,000 in the '
        'prior quarter, reflecting stronger enterprise-tier adoption.'
    )

    # --- Section 2: Competitive Landscape ---
    doc.add_heading('Competitive Landscape', level=1)

    doc.add_paragraph(
        'The competitive environment remains dynamic, with three notable shifts during '
        'Q1. First, Meridian Technologies acquired DataSync Corp for $2.3 billion, '
        'consolidating their position in the data integration space. Second, NovaPeak '
        'Software launched a direct competitor to our core analytics module, pricing it '
        'at a 20% discount to our standard tier.'
    )
    doc.add_paragraph(
        'Despite increased competition, our net promoter score improved to 67 (from 61 '
        'in Q4 2024), and customer retention rate held steady at 94.3%. Win rates against '
        'key competitors were: Meridian Technologies 58%, NovaPeak Software 71%, and '
        'Orion Systems 44%. The lower win rate against Orion reflects their strength in '
        'the government vertical, where procurement cycles favor established incumbents.'
    )
    doc.add_paragraph(
        'Product differentiation remains our strongest competitive advantage. Our '
        'proprietary real-time collaboration engine and cross-platform compatibility '
        'were cited as top purchase factors by 82% of new enterprise customers surveyed '
        'during the quarter.'
    )

    # --- Section 3: Growth Strategy ---
    doc.add_heading('Growth Strategy', level=1)

    doc.add_paragraph(
        'Our growth strategy for the remainder of 2025 focuses on three pillars: '
        'geographic expansion, product innovation, and strategic partnerships. We plan '
        'to establish regional offices in Singapore and Frankfurt by Q3, targeting '
        'underserved mid-market segments in Southeast Asia and Central Europe.'
    )
    doc.add_paragraph(
        'On the product side, our roadmap includes the launch of an AI-powered '
        'document intelligence module in Q2 and a revamped mobile experience in Q3. '
        'Early beta testing with 150 customers showed a 34% improvement in document '
        'processing speed and a 22% reduction in manual review steps.'
    )
    doc.add_paragraph(
        'Strategic partnerships with Fujitsu, SAP, and Salesforce are expected to '
        'generate $18 million in co-sell revenue by year-end. The Fujitsu partnership '
        'alone is projected to open access to over 400 enterprise accounts in Japan '
        'and South Korea that we currently cannot reach through direct sales channels.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
