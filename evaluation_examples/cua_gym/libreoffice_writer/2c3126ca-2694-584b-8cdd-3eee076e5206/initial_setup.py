"""
Initial Setup: Architecture document with numbered figure captions.
Task ID: writer_tm_087
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_087'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bookmark(para, bookmark_name, text):
    """Add a bookmark around text in a paragraph."""
    import random
    bm_id = str(random.randint(1000, 99999))

    # Bookmark start
    bm_start = para._element.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): bm_id,
        qn('w:name'): bookmark_name,
    })

    # The run with text
    run = para.add_run(text)

    # Bookmark end
    bm_end = para._element.makeelement(qn('w:bookmarkEnd'), {
        qn('w:id'): bm_id,
    })

    # Insert bookmark start before the run, end after
    run_elem = run._element
    run_elem.addprevious(bm_start)
    run_elem.addnext(bm_end)

    return run


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title Page ---
    title = doc.add_heading("Enterprise Network Architecture", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Technical Reference Document")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run("Version 3.2 — March 2025\nPrepared by: Infrastructure Engineering Team")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # === PAGE 2: Introduction ===
    doc.add_heading("1. Introduction", level=1)

    doc.add_paragraph(
        "This document provides a comprehensive overview of the enterprise network "
        "architecture deployed across all regional offices. It covers the physical and "
        "logical topology, security zones, redundancy mechanisms, and integration points "
        "with cloud services."
    )

    doc.add_paragraph(
        "The architecture was designed to support high availability with 99.99% uptime SLA, "
        "scalable bandwidth provisioning up to 40 Gbps per site, and zero-trust security "
        "principles aligned with NIST SP 800-207 guidelines."
    )

    # Figure 1: High-Level Overview
    fig1_para = doc.add_paragraph()
    fig1_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig1_para.add_run("[Diagram: High-Level Enterprise Network Overview]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    cap1 = doc.add_paragraph()
    cap1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_bookmark(cap1, "_Ref_Fig1", "Figure 1: High-Level Enterprise Network Overview")
    for run in cap1.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_paragraph(
        "As shown in the diagram above, the enterprise network spans three primary layers: "
        "the core backbone, the distribution layer, and the access layer. Each layer is "
        "designed with redundant paths to eliminate single points of failure."
    )

    # KEY paragraph - cursor location: "The network layout is illustrated in "
    ref_para = doc.add_paragraph(
        "The network layout is illustrated in "
    )
    # NOTE: No cross-reference here in initial state. The agent must insert one.

    doc.add_paragraph(
        "The following sections describe each component in detail, including hardware "
        "specifications, software configurations, and integration protocols used throughout "
        "the deployment."
    )

    doc.add_page_break()

    # === PAGE 3: Core Network ===
    doc.add_heading("2. Core Network Infrastructure", level=1)

    doc.add_paragraph(
        "The core network backbone consists of dual Cisco Nexus 9500 chassis switches "
        "deployed in a VXLAN-EVPN fabric. This provides Layer 2 and Layer 3 connectivity "
        "across all data center pods with sub-millisecond failover capabilities."
    )

    doc.add_paragraph(
        "BGP is used as the underlay routing protocol with OSPF for intra-pod communication. "
        "The VXLAN overlay supports up to 16 million virtual network segments, allowing "
        "fine-grained microsegmentation of workloads."
    )

    # Figure 2
    fig2_para = doc.add_paragraph()
    fig2_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig2_para.add_run("[Diagram: Core Switch Fabric Topology]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    cap2 = doc.add_paragraph()
    cap2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_bookmark(cap2, "_Ref_Fig2", "Figure 2: Core Switch Fabric Topology")
    for run in cap2.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_paragraph(
        "Each spine switch maintains a full mesh of eBGP sessions with all leaf switches. "
        "The fabric supports 100GbE uplinks with 25GbE host-facing ports, delivering an "
        "aggregate throughput of 12.8 Tbps per pod."
    )

    doc.add_page_break()

    # === PAGE 4: Distribution Layer ===
    doc.add_heading("3. Distribution Layer", level=1)

    doc.add_paragraph(
        "The distribution layer aggregates traffic from multiple access switches and enforces "
        "policy-based routing decisions. Palo Alto PA-5260 firewalls are deployed at each "
        "distribution node to provide east-west traffic inspection."
    )

    doc.add_paragraph(
        "Quality of Service policies are applied at this layer to prioritize voice and video "
        "traffic. DSCP markings are preserved end-to-end, with eight traffic classes mapped "
        "to four hardware queues."
    )

    # Figure 3
    fig3_para = doc.add_paragraph()
    fig3_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig3_para.add_run("[Diagram: Distribution Layer with Firewall Zones]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    cap3 = doc.add_paragraph()
    cap3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_bookmark(cap3, "_Ref_Fig3", "Figure 3: Distribution Layer with Firewall Zones")
    for run in cap3.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_paragraph(
        "Inter-VLAN routing is handled by the distribution switches using virtual routing "
        "instances (VRFs). Each tenant or department is assigned a dedicated VRF to maintain "
        "traffic isolation while sharing the physical infrastructure."
    )

    doc.add_paragraph(
        "Redundancy is achieved through Virtual Router Redundancy Protocol (VRRP) with "
        "sub-second failover. Primary and backup distribution switches synchronize state "
        "tables via a dedicated management VLAN to ensure seamless transition during failures."
    )

    doc.add_page_break()

    # === PAGE 5: Access Layer ===
    doc.add_heading("4. Access Layer", level=1)

    doc.add_paragraph(
        "The access layer provides end-user connectivity through Cisco Catalyst 9300 switches "
        "and Aruba AP-535 wireless access points. Each floor of every building is served by "
        "a dedicated access switch stack with redundant uplinks to the distribution layer."
    )

    doc.add_paragraph(
        "802.1X port-based authentication is enforced on all wired ports, integrating with "
        "Cisco ISE for dynamic VLAN assignment based on user role and device posture. "
        "Wireless authentication supports both WPA3-Enterprise and certificate-based EAP-TLS."
    )

    # Figure 4
    fig4_para = doc.add_paragraph()
    fig4_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig4_para.add_run("[Diagram: Access Layer Floor Plan Layout]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    cap4 = doc.add_paragraph()
    cap4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_bookmark(cap4, "_Ref_Fig4", "Figure 4: Access Layer Floor Plan Layout")
    for run in cap4.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_paragraph(
        "Power over Ethernet (PoE++) is available on all access ports, supporting up to 60W "
        "per port for IP phones, cameras, and wireless access points. The switch stacks "
        "support IEEE 802.3bt with a total PoE budget of 1440W per stack member."
    )

    doc.add_page_break()

    # === PAGE 6: Security Architecture ===
    doc.add_heading("5. Security Architecture", level=1)

    doc.add_paragraph(
        "The security architecture follows a defense-in-depth strategy with multiple layers "
        "of protection. The perimeter is secured by Fortinet FortiGate 3600E next-generation "
        "firewalls configured in active-active high availability mode."
    )

    doc.add_paragraph(
        "Intrusion Detection and Prevention Systems (IDS/IPS) are deployed at strategic "
        "points within the network. Signature-based and behavioral analysis engines work in "
        "tandem to identify and block known threats as well as zero-day attacks."
    )

    doc.add_paragraph(
        "Network Access Control (NAC) ensures that only compliant devices are granted access. "
        "Devices failing the posture assessment are quarantined in a remediation VLAN with "
        "limited internet access for patching and updates."
    )

    doc.add_page_break()

    # === PAGE 7: Cloud Integration ===
    doc.add_heading("6. Cloud Integration", level=1)

    doc.add_paragraph(
        "Hybrid cloud connectivity is established through AWS Direct Connect and Azure "
        "ExpressRoute circuits. Each circuit provides 10 Gbps of dedicated bandwidth with "
        "automatic failover to IPsec VPN tunnels over the public internet."
    )

    doc.add_paragraph(
        "SD-WAN overlay networking is used to optimize traffic routing between on-premises "
        "data centers and cloud workloads. Application-aware routing policies direct traffic "
        "based on real-time path quality metrics including latency, jitter, and packet loss."
    )

    doc.add_paragraph(
        "Cloud security groups and network ACLs mirror the on-premises segmentation policy. "
        "A centralized management plane provides unified visibility across both environments "
        "through Cisco ACI Multi-Site and Azure Arc integration."
    )

    doc.add_page_break()

    # === PAGE 8: Monitoring & Management ===
    doc.add_heading("7. Monitoring and Management", level=1)

    doc.add_paragraph(
        "Network monitoring is performed by a suite of tools including SolarWinds NPM for "
        "device-level metrics, Splunk for log aggregation and correlation, and ThousandEyes "
        "for synthetic monitoring of end-to-end application paths."
    )

    doc.add_paragraph(
        "SNMP v3 is used for device polling with 60-second intervals for critical devices "
        "and 300-second intervals for standard equipment. Syslog messages are forwarded in "
        "real-time to the Splunk cluster over TLS-encrypted connections."
    )

    doc.add_paragraph(
        "Automated remediation workflows are configured through Ansible Tower, enabling "
        "self-healing for common issues such as port flaps, BGP session resets, and "
        "certificate expirations. The ITSM integration with ServiceNow ensures all automated "
        "actions are tracked with proper change records."
    )

    doc.add_page_break()

    # === PAGE 9: Network Architecture Diagram (Figure 5) ===
    doc.add_heading("8. Complete Network Architecture", level=1)

    doc.add_paragraph(
        "The following diagram presents the complete end-to-end network architecture, "
        "consolidating all layers, security zones, and cloud integration points discussed "
        "in the preceding sections."
    )

    # Figure 5 - the target of the cross-reference
    fig5_para = doc.add_paragraph()
    fig5_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fig5_para.add_run("[Diagram: Complete Network Architecture - All Layers]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    cap5 = doc.add_paragraph()
    cap5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_bookmark(cap5, "_Ref_Fig5", "Figure 5: Network Architecture Diagram")
    for run in cap5.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_paragraph(
        "This comprehensive view shows the interconnections between the core VXLAN-EVPN "
        "fabric, distribution firewalls, access layer switches, and cloud on-ramp circuits. "
        "The color coding follows the convention established in the security zone matrix: "
        "green for trusted zones, yellow for semi-trusted, and red for untrusted zones."
    )

    doc.add_paragraph(
        "Key performance indicators for the overall architecture include an aggregate "
        "north-south throughput of 80 Gbps, east-west capacity of 2.4 Tbps, and a maximum "
        "convergence time of 250 milliseconds for any single component failure."
    )

    doc.add_page_break()

    # === PAGE 10: Appendix ===
    doc.add_heading("Appendix A: Hardware Inventory", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Device", "Model", "Quantity", "Location"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    inventory = [
        ["Core Switch", "Cisco Nexus 9508", "4", "DC-East, DC-West"],
        ["Distribution FW", "Palo Alto PA-5260", "8", "All Sites"],
        ["Access Switch", "Cisco Catalyst 9300", "120", "All Floors"],
        ["Wireless AP", "Aruba AP-535", "450", "All Floors"],
        ["WAN Router", "Cisco ASR 1001-HX", "6", "DC-East, DC-West, HQ"],
        ["Load Balancer", "F5 BIG-IP i5800", "4", "DC-East, DC-West"],
        ["DNS/DHCP", "Infoblox NIOS 8.5", "2", "DC-East, DC-West"],
        ["NTP Server", "Meinberg LANTIME M300", "2", "DC-East, DC-West"],
    ]
    for row_data in inventory:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
