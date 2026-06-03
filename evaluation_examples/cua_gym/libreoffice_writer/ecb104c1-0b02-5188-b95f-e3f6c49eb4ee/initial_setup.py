"""
Initial Setup: Create a Writer document with a 5-column, 11-row data table (all white background)
Task ID: writer_biz_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_053'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Q1 2025 Sales Performance Report', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add an intro paragraph
    intro = doc.add_paragraph(
        'The following table summarizes the quarterly sales performance '
        'for each regional representative. All figures are reported in USD.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Create 5-column, 11-row table (1 header + 10 data rows)
    table = doc.add_table(rows=11, cols=5)
    table.style = 'Table Grid'

    # Header row
    headers = ['Representative', 'Region', 'Q1 Revenue', 'Units Sold', 'Target Met']
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # 10 data rows with realistic business data
    data = [
        ['Sarah Chen',       'Northeast',  '$87,450',   '342',  'Yes'],
        ['Marcus Johnson',   'Southeast',  '$63,210',   '258',  'No'],
        ['Elena Rodriguez',  'West Coast', '$91,875',   '367',  'Yes'],
        ['David Kim',        'Midwest',    '$54,320',   '215',  'No'],
        ['Priya Patel',      'Southwest',  '$72,680',   '291',  'Yes'],
        ['James O\'Brien',   'Northwest',  '$68,940',   '276',  'Yes'],
        ['Aisha Williams',   'Central',    '$59,130',   '234',  'No'],
        ['Robert Fischer',   'Mid-Atlantic','$83,260',  '330',  'Yes'],
        ['Mei-Lin Tang',     'Pacific',    '$76,540',   '305',  'Yes'],
        ['Carlos Herrera',   'Mountain',   '$61,890',   '247',  'No'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            run = cell.paragraphs[0].add_run(value)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # Add a closing paragraph
    doc.add_paragraph('')
    closing = doc.add_paragraph(
        'Note: Representatives who met their quarterly targets are eligible '
        'for the performance bonus program effective April 2025.'
    )
    closing.runs[0].font.size = Pt(10)
    closing.runs[0].italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
