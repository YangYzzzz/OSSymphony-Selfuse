"""
Reward Script: Apply alternating row shading to a 10-row data table
Task ID: writer_biz_053
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Even data rows (2,4,6,8,10) have #F2F2F2 shading
  Component 2 (0.3): Alternating pattern is correct — even rows shaded AND odd rows unshaded
  Component 3 (0.2): Complete alternating pattern including header unshaded
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_053'


def get_cell_fill(cell):
    """Extract the fill/shading color from a table cell. Returns uppercase hex or None."""
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    if fill is None or fill.lower() in ('auto', 'none', 'ffffff'):
        return None
    return fill.upper()


def is_gray_shading(fill_val):
    """Check if the fill value is approximately #F2F2F2 (light gray)."""
    if fill_val is None:
        return False
    target = 0xF2
    try:
        r = int(fill_val[0:2], 16)
        g = int(fill_val[2:4], 16)
        b = int(fill_val[4:6], 16)
        if abs(r - target) <= 5 and abs(g - target) <= 5 and abs(b - target) <= 5:
            return True
    except (ValueError, IndexError):
        pass
    return False


def is_unshaded(fill_val):
    """Check if a cell has no shading (white or no fill)."""
    return fill_val is None


def verify_task(file_path):
    """
    Verify alternating row shading in the data table.
    Returns: float between 0.0 and 1.0

    Key design: ALL components are anchored to the task-introduced change
    (even data rows getting F2F2F2 shading). Components 2 and 3 are compound
    checks that require BOTH the shading on even rows AND the absence on odd rows.
    This ensures initial_env (no shading anywhere) scores 0.0.
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table with 11 rows
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    if len(table.rows) < 11:
        print(f"FAIL: Table has {len(table.rows)} rows, expected at least 11")
        print("REWARD: 0.0")
        return 0.0

    # Gather shading info for all rows
    even_data_rows = [2, 4, 6, 8, 10]  # 0-indexed table rows
    odd_data_rows = [1, 3, 5, 7, 9]

    # Count correctly shaded even-row cells
    even_shaded_count = 0
    even_total = 0
    for row_idx in even_data_rows:
        for cell in table.rows[row_idx].cells:
            even_total += 1
            if is_gray_shading(get_cell_fill(cell)):
                even_shaded_count += 1

    # Count correctly unshaded odd-row cells
    odd_unshaded_count = 0
    odd_total = 0
    for row_idx in odd_data_rows:
        for cell in table.rows[row_idx].cells:
            odd_total += 1
            if is_unshaded(get_cell_fill(cell)):
                odd_unshaded_count += 1

    # Count correctly unshaded header cells
    header_unshaded_count = 0
    header_total = 0
    for cell in table.rows[0].cells:
        header_total += 1
        if is_unshaded(get_cell_fill(cell)):
            header_unshaded_count += 1

    all_even_shaded = (even_total > 0 and even_shaded_count == even_total)
    all_odd_unshaded = (odd_total > 0 and odd_unshaded_count == odd_total)
    all_header_unshaded = (header_total > 0 and header_unshaded_count == header_total)

    # Component 1: Even data rows have #F2F2F2 shading (0.5 points)
    # This is the PRIMARY task change — must fail on initial (no shading)
    try:
        if all_even_shaded:
            print(f"PASS: Component 1 — All {even_total} cells in even data rows have F2F2F2 shading (0.5 pts)")
            total_score += 0.5
        elif even_total > 0:
            partial = 0.5 * (even_shaded_count / even_total)
            print(f"PARTIAL: Component 1 — {even_shaded_count}/{even_total} even-row cells shaded correctly ({partial:.2f} pts)")
            total_score += partial
        else:
            print("FAIL: Component 1 — No cells to check in even data rows")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Alternating pattern — even rows shaded AND odd rows unshaded (0.3 points)
    # Compound check: requires even rows to be shaded (task change) as gate condition
    # This prevents scoring on initial_env where odd rows are also unshaded but even rows aren't shaded
    try:
        if all_even_shaded and all_odd_unshaded:
            print(f"PASS: Component 2 — Alternating pattern correct: even shaded + odd unshaded (0.3 pts)")
            total_score += 0.3
        elif all_even_shaded and odd_total > 0:
            partial = 0.3 * (odd_unshaded_count / odd_total)
            print(f"PARTIAL: Component 2 — Even rows shaded but {odd_total - odd_unshaded_count} odd-row cells also shaded ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Even rows not fully shaded (gate condition failed): {even_shaded_count}/{even_total}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Complete pattern — even rows shaded AND header unshaded (0.2 points)
    # Compound check: requires even rows to be shaded (task change) as gate condition
    try:
        if all_even_shaded and all_header_unshaded:
            print(f"PASS: Component 3 — Header unshaded with even rows correctly shaded (0.2 pts)")
            total_score += 0.2
        elif all_even_shaded and header_total > 0:
            partial = 0.2 * (header_unshaded_count / header_total)
            print(f"PARTIAL: Component 3 — Even rows shaded but {header_total - header_unshaded_count} header cells also shaded ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Even rows not fully shaded (gate condition failed): {even_shaded_count}/{even_total}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
