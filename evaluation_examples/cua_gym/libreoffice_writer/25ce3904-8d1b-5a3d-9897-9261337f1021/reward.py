"""
Reward Script: Three-level outline numbering for commercial lease contract
Task ID: writer_legal_045
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Heading styles applied (Heading 1/2/3) to article/section/subsection paras
  Component 2 (0.30): Outline numbering (numPr) linked to heading paragraphs
  Component 3 (0.40): Numbering format definition matches legal convention
    - Level 0: upperRoman, 'ARTICLE %1'
    - Level 1: decimal, 'Section %1.%2'
    - Level 2: lowerLetter, '(%3)'
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_045'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    w_ns = ns['w']

    # =========================================================================
    # Component 1: Heading styles applied to article/section/subsection (0.30)
    # In initial doc, ALL paragraphs are Normal. In golden, 58 use Heading 1/2/3.
    # =========================================================================
    try:
        h1_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 1')
        h2_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 2')
        h3_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 3')
        total_headings = h1_count + h2_count + h3_count

        print(f"  Heading 1 count: {h1_count} (expected ~10)")
        print(f"  Heading 2 count: {h2_count} (expected ~23)")
        print(f"  Heading 3 count: {h3_count} (expected ~25)")

        # Need at least some of each heading level, and total should be substantial
        if h1_count >= 8 and h2_count >= 15 and h3_count >= 15:
            print(f"PASS: Component 1 — Heading styles applied: {total_headings} heading paragraphs (0.30 pts)")
            total_score += 0.30
        elif h1_count >= 5 and h2_count >= 10 and h3_count >= 10:
            print(f"PARTIAL: Component 1 — Some heading styles applied: H1={h1_count}, H2={h2_count}, H3={h3_count} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Insufficient heading styles: H1={h1_count}, H2={h2_count}, H3={h3_count}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Outline numbering (numPr) linked to heading paragraphs (0.30)
    # In initial doc, no paragraphs have numPr. In golden, all 58 headings have numPr.
    # =========================================================================
    try:
        heading_with_numpr = 0
        heading_total = 0

        for para in doc.paragraphs:
            if para.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
                heading_total += 1
                pPr = para._element.find('w:pPr', ns)
                if pPr is not None:
                    numPr = pPr.find('w:numPr', ns)
                    if numPr is not None:
                        heading_with_numpr += 1

        print(f"  Headings with numPr: {heading_with_numpr}/{heading_total}")

        if heading_total > 0 and heading_with_numpr >= heading_total * 0.8:
            print(f"PASS: Component 2 — Outline numbering linked to headings ({heading_with_numpr}/{heading_total}) (0.30 pts)")
            total_score += 0.30
        elif heading_total > 0 and heading_with_numpr >= heading_total * 0.5:
            print(f"PARTIAL: Component 2 — Some numbering linked ({heading_with_numpr}/{heading_total}) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Numbering not linked to headings ({heading_with_numpr}/{heading_total})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: Numbering format definition matches legal convention (0.40)
    # Check the abstractNum definition for the correct level formats:
    #   Level 0: upperRoman, 'ARTICLE %1'
    #   Level 1: decimal, 'Section %1.%2'
    #   Level 2: lowerLetter, '(%3)'
    # In initial doc, there is no such numbering definition at all.
    # =========================================================================
    try:
        # Find any numId used by heading paragraphs
        heading_numids = set()
        for para in doc.paragraphs:
            if para.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
                pPr = para._element.find('w:pPr', ns)
                if pPr is not None:
                    numPr = pPr.find('w:numPr', ns)
                    if numPr is not None:
                        numId_el = numPr.find('w:numId', ns)
                        if numId_el is not None:
                            heading_numids.add(numId_el.get(f'{{{w_ns}}}val'))

        if not heading_numids:
            print("FAIL: Component 3 — No numId found on heading paragraphs")
        else:
            # Get the numbering part and find the abstractNum
            numbering_part = doc.part.numbering_part
            numbering_elem = numbering_part.element

            level_checks = {0: False, 1: False, 2: False}
            level_details = {}

            for numId_val in heading_numids:
                # Find num element
                for num_el in numbering_elem.findall('.//w:num', ns):
                    if num_el.get(f'{{{w_ns}}}numId') == numId_val:
                        absNumId_el = num_el.find('w:abstractNumId', ns)
                        if absNumId_el is None:
                            continue
                        absNumId = absNumId_el.get(f'{{{w_ns}}}val')

                        # Find abstractNum
                        for absNum in numbering_elem.findall('.//w:abstractNum', ns):
                            if absNum.get(f'{{{w_ns}}}abstractNumId') == absNumId:
                                for lvl in absNum.findall('w:lvl', ns):
                                    ilvl = int(lvl.get(f'{{{w_ns}}}ilvl', '-1'))
                                    numFmt_el = lvl.find('w:numFmt', ns)
                                    lvlText_el = lvl.find('w:lvlText', ns)
                                    numFmt = numFmt_el.get(f'{{{w_ns}}}val') if numFmt_el is not None else None
                                    lvlText = lvlText_el.get(f'{{{w_ns}}}val') if lvlText_el is not None else None

                                    level_details[ilvl] = (numFmt, lvlText)

                                    # Level 0: upperRoman, contains 'ARTICLE'
                                    if ilvl == 0:
                                        if numFmt == 'upperRoman' and lvlText and 'ARTICLE' in lvlText.upper():
                                            level_checks[0] = True

                                    # Level 1: decimal, contains 'Section'
                                    elif ilvl == 1:
                                        if numFmt == 'decimal' and lvlText and 'ection' in lvlText:
                                            level_checks[1] = True

                                    # Level 2: lowerLetter, pattern like '(%3)' or similar
                                    elif ilvl == 2:
                                        if numFmt == 'lowerLetter' and lvlText and '(' in lvlText and ')' in lvlText:
                                            level_checks[2] = True

            for lvl_num, (fmt, txt) in sorted(level_details.items()):
                if lvl_num <= 2:
                    status = "OK" if level_checks.get(lvl_num) else "MISMATCH"
                    print(f"  Level {lvl_num}: numFmt={fmt}, lvlText={txt!r} [{status}]")

            passed = sum(1 for v in level_checks.values() if v)
            if passed == 3:
                print(f"PASS: Component 3 — All 3 numbering levels match legal convention (0.40 pts)")
                total_score += 0.40
            elif passed == 2:
                print(f"PARTIAL: Component 3 — {passed}/3 levels correct (0.27 pts)")
                total_score += 0.27
            elif passed == 1:
                print(f"PARTIAL: Component 3 — {passed}/3 levels correct (0.13 pts)")
                total_score += 0.13
            else:
                print(f"FAIL: Component 3 — No levels match legal convention")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
