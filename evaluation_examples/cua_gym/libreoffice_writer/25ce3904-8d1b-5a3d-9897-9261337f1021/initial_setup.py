"""
Initial Setup: Commercial lease document with manual numbering
Task ID: writer_legal_045
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('COMMERCIAL LEASE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Preamble
    preamble = doc.add_paragraph()
    preamble.add_run(
        'This Commercial Lease Agreement ("Agreement") is entered into as of '
        'March 15, 2025, by and between Meridian Properties LLC ("Landlord") '
        'and Cascade Technology Solutions Inc. ("Tenant").'
    )
    preamble.paragraph_format.space_after = Pt(12)

    # ---- Define the lease content with manual numbering ----
    # Each entry: (level, text)
    # level 1 = ARTICLE, level 2 = Section, level 3 = subsection
    lease_content = [
        # ARTICLE I
        (1, "ARTICLE I - PREMISES AND TERM"),
        (2, "Section 1.01 - Premises"),
        (0, "Landlord hereby leases to Tenant approximately 12,500 square feet of office space located at 4200 Westfield Boulevard, Suite 300, Portland, Oregon 97201 (the \"Premises\"), as outlined in Exhibit A attached hereto."),
        (2, "Section 1.02 - Term"),
        (0, "The initial term of this Lease shall commence on May 1, 2025 (the \"Commencement Date\") and shall expire on April 30, 2030 (the \"Expiration Date\"), unless sooner terminated in accordance with the provisions hereof."),
        (3, "(a) Tenant shall have two (2) options to renew the Lease for additional five-year periods upon written notice to Landlord no later than one hundred eighty (180) days prior to the Expiration Date."),
        (3, "(b) Renewal terms shall be at the then-prevailing market rate, not to exceed 105% of the Base Rent for the preceding year."),
        (2, "Section 1.03 - Early Access"),
        (0, "Tenant shall have access to the Premises thirty (30) days prior to the Commencement Date for the sole purpose of installing fixtures, equipment, and telecommunications infrastructure."),

        # ARTICLE II
        (1, "ARTICLE II - RENT AND ADDITIONAL CHARGES"),
        (2, "Section 2.01 - Base Rent"),
        (0, "Tenant shall pay to Landlord base rent in the amount of $31,250.00 per month ($375,000.00 annually), payable in advance on the first day of each calendar month during the Term."),
        (3, "(a) Base Rent shall increase by three percent (3%) annually on each anniversary of the Commencement Date."),
        (3, "(b) Rent for any partial month shall be prorated on a per diem basis."),
        (3, "(c) All payments shall be made by electronic funds transfer to Landlord's designated account."),
        (2, "Section 2.02 - Security Deposit"),
        (0, "Tenant shall deposit with Landlord the sum of $62,500.00 as a security deposit upon execution of this Agreement."),
        (2, "Section 2.03 - Operating Expenses"),
        (0, "Tenant shall pay its proportionate share (currently 18.7%) of all operating expenses for the Building in excess of the Base Year (calendar year 2025) expenses."),

        # ARTICLE III
        (1, "ARTICLE III - USE OF PREMISES"),
        (2, "Section 3.01 - Permitted Use"),
        (0, "The Premises shall be used and occupied by Tenant solely for general office purposes, software development, and ancillary activities consistent with a Class A office environment."),
        (2, "Section 3.02 - Prohibited Uses"),
        (3, "(a) Tenant shall not use the Premises for any purpose that violates applicable zoning ordinances or federal, state, or local law."),
        (3, "(b) Tenant shall not conduct any activity that creates excessive noise, vibration, or odors that disturb other tenants."),
        (3, "(c) Tenant shall not store hazardous materials on the Premises except in quantities customary for office use."),

        # ARTICLE IV
        (1, "ARTICLE IV - MAINTENANCE AND REPAIRS"),
        (2, "Section 4.01 - Landlord Obligations"),
        (0, "Landlord shall maintain the structural components of the Building, including the roof, exterior walls, foundation, and all common area mechanical, electrical, plumbing, and HVAC systems."),
        (2, "Section 4.02 - Tenant Obligations"),
        (0, "Tenant shall maintain the interior of the Premises in good condition, including all non-structural repairs, interior painting, carpet replacement, and maintenance of Tenant-installed equipment."),
        (3, "(a) Tenant shall promptly notify Landlord of any condition requiring repair that is Landlord's responsibility."),
        (3, "(b) If Tenant fails to maintain the Premises, Landlord may perform necessary repairs and charge Tenant for reasonable costs incurred."),

        # ARTICLE V
        (1, "ARTICLE V - ALTERATIONS AND IMPROVEMENTS"),
        (2, "Section 5.01 - Tenant Improvements"),
        (0, "Landlord shall provide a Tenant Improvement Allowance of $50.00 per rentable square foot ($625,000.00 total) for the build-out of the Premises according to plans approved by Landlord."),
        (2, "Section 5.02 - Subsequent Alterations"),
        (0, "Tenant may make non-structural alterations costing less than $25,000.00 without Landlord's prior written consent. All other alterations require Landlord's written approval, which shall not be unreasonably withheld."),
        (3, "(a) All alterations shall comply with applicable building codes and ADA requirements."),
        (3, "(b) Tenant shall remove all alterations and restore the Premises at the end of the Term, unless Landlord elects to retain them."),

        # ARTICLE VI
        (1, "ARTICLE VI - INSURANCE AND INDEMNIFICATION"),
        (2, "Section 6.01 - Tenant Insurance"),
        (0, "Tenant shall maintain commercial general liability insurance with limits of not less than $2,000,000.00 per occurrence and $5,000,000.00 aggregate, naming Landlord as an additional insured."),
        (2, "Section 6.02 - Property Insurance"),
        (0, "Tenant shall maintain insurance covering all of Tenant's personal property, trade fixtures, and improvements within the Premises at full replacement cost."),
        (3, "(a) All insurance policies shall be issued by carriers rated A- or better by A.M. Best Company."),
        (3, "(b) Certificates of insurance shall be delivered to Landlord annually and upon any policy renewal."),

        # ARTICLE VII
        (1, "ARTICLE VII - ASSIGNMENT AND SUBLETTING"),
        (2, "Section 7.01 - Consent Required"),
        (0, "Tenant shall not assign this Lease or sublet all or any portion of the Premises without the prior written consent of Landlord, which shall not be unreasonably withheld, conditioned, or delayed."),
        (2, "Section 7.02 - Conditions of Transfer"),
        (3, "(a) Any approved assignee or subtenant shall assume all obligations of Tenant under this Lease."),
        (3, "(b) Tenant shall remain primarily liable for all Lease obligations notwithstanding any approved assignment."),
        (3, "(c) Any assignment or sublease profits exceeding Tenant's rent obligation shall be shared equally between Landlord and Tenant."),

        # ARTICLE VIII
        (1, "ARTICLE VIII - DEFAULT AND REMEDIES"),
        (2, "Section 8.01 - Events of Default"),
        (0, "The following shall constitute events of default under this Lease:"),
        (3, "(a) Failure to pay rent within ten (10) business days after written notice of nonpayment."),
        (3, "(b) Failure to perform any non-monetary obligation within thirty (30) days after written notice."),
        (3, "(c) Filing of a petition in bankruptcy or insolvency by or against Tenant."),
        (2, "Section 8.02 - Landlord Remedies"),
        (0, "Upon the occurrence of an event of default, Landlord may terminate this Lease, re-enter the Premises, accelerate all rent due for the remainder of the Term, and pursue all other remedies available at law or in equity."),

        # ARTICLE IX
        (1, "ARTICLE IX - DAMAGE AND DESTRUCTION"),
        (2, "Section 9.01 - Partial Damage"),
        (0, "If the Premises are partially damaged by fire or other casualty, Landlord shall repair the damage within one hundred eighty (180) days, and rent shall abate proportionally during the repair period."),
        (2, "Section 9.02 - Total Destruction"),
        (0, "If the Premises are substantially destroyed (more than 50% of the rentable area), either party may terminate this Lease by written notice within sixty (60) days of the casualty."),
        (3, "(a) In the event of termination, Tenant's rent obligation shall cease as of the date of the casualty."),
        (3, "(b) Landlord shall refund any prepaid rent and the security deposit within thirty (30) days of termination."),

        # ARTICLE X
        (1, "ARTICLE X - MISCELLANEOUS PROVISIONS"),
        (2, "Section 10.01 - Governing Law"),
        (0, "This Agreement shall be governed by and construed in accordance with the laws of the State of Oregon."),
        (2, "Section 10.02 - Entire Agreement"),
        (0, "This Lease constitutes the entire agreement between the parties and supersedes all prior negotiations, representations, and agreements, whether written or oral."),
        (2, "Section 10.03 - Notices"),
        (3, "(a) All notices shall be in writing and delivered by certified mail, return receipt requested, or by nationally recognized overnight courier."),
        (3, "(b) Notices to Landlord: Meridian Properties LLC, 1700 Pioneer Court, Suite 800, Portland, OR 97204."),
        (3, "(c) Notices to Tenant: Cascade Technology Solutions Inc., 4200 Westfield Boulevard, Suite 300, Portland, OR 97201."),
    ]

    # Write the lease content with manual numbering (no heading styles)
    for level, text in lease_content:
        if level == 1:
            # Article headings - bold, larger font, but NO Heading style
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
        elif level == 2:
            # Section headings - bold, but NO Heading style
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
        elif level == 3:
            # Subsection items - indented, with manual (a), (b) numbering
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.75)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(text)
            run.font.size = Pt(11)
        else:
            # Body text
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            run.font.size = Pt(11)

    # Signature block
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(24)
    sig.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.').bold = True

    for party in [('LANDLORD:', 'Meridian Properties LLC', 'Victoria Hartwell, Managing Partner'),
                  ('TENANT:', 'Cascade Technology Solutions Inc.', 'David Nakamura, Chief Executive Officer')]:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.add_run(party[0]).bold = True
        p2 = doc.add_paragraph()
        p2.add_run(party[1])
        p3 = doc.add_paragraph()
        p3.add_run('By: ____________________________')
        p4 = doc.add_paragraph()
        p4.add_run(f'Name: {party[2]}')
        p5 = doc.add_paragraph()
        p5.add_run('Date: ____________________________')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
