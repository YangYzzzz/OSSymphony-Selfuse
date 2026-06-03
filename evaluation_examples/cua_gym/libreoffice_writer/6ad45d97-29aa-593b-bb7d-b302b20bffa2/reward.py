"""
Reward Script: Insert bookmark 'appendix_start' at 'Appendix A: Raw Data' heading and update TOC to 6 entries
Task ID: writer_struct_071
Domain: libreoffice_writer
Scoring:
  Component 1: Bookmark 'appendix_start' exists at 'Appendix A: Raw Data' heading (0.5 pts)
  Component 2: TOC has 6 entries (was 5 before task) (0.3 pts)
  Component 3: TOC includes an entry containing 'Appendix A: Raw Data' with a page number (0.2 pts)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_071'
FILE_PATH = os.path.join(WORKDIR, 'survey_analysis.docx')

# XML namespace constants
W_BOOKMARK_START = qn('w:bookmarkStart')
W_NAME = qn('w:name')
W_T = qn('w:t')


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Insert bookmark 'appendix_start' at 'Appendix A: Raw Data' heading
          and update TOC to include the appendix entry (6 entries total).
    """
    total_score = 0.0

    # Load document — if this fails, return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Bookmark 'appendix_start' exists and is placed at
    #   the 'Appendix A: Raw Data' Heading 1 paragraph (0.5 points)
    #   This verifies the core task requirement: the bookmark must exist
    #   at the correct heading location.
    try:
        body = doc.element.body
        all_bookmarks = body.findall('.//' + W_BOOKMARK_START)

        # Find if a bookmark named 'appendix_start' exists
        appendix_bookmark = None
        for bm in all_bookmarks:
            if bm.get(W_NAME) == 'appendix_start':
                appendix_bookmark = bm
                break

        if appendix_bookmark is None:
            print("FAIL: Component 1 — bookmark 'appendix_start' not found in document")
        else:
            # Verify that this bookmark is inside the 'Appendix A: Raw Data' heading paragraph
            parent_para = appendix_bookmark.getparent()
            parent_text = ''
            if parent_para is not None:
                for t in parent_para.iter(W_T):
                    parent_text += t.text or ''

            if 'Appendix A' in parent_text and 'Raw Data' in parent_text:
                print(f"PASS: Component 1 — bookmark 'appendix_start' found at heading "
                      f"'Appendix A: Raw Data' (parent text: {parent_text!r}) (0.5 pts)")
                total_score += 0.5
            else:
                # Bookmark exists but may not be at the correct paragraph
                # Still give partial credit for bookmark existence
                print(f"PARTIAL: Component 1 — bookmark 'appendix_start' exists but parent "
                      f"paragraph text is {parent_text!r}, expected 'Appendix A: Raw Data'. "
                      f"Awarding 0.2 pts for bookmark existence only.")
                total_score += 0.2
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: TOC has 6 entries (was 5 initially, must now include appendix) (0.3 points)
    #   The original TOC had 5 entries. After the task, it should have 6 entries.
    try:
        toc_entries = []
        in_toc = False
        for para in doc.paragraphs:
            if para.style.name == 'Heading 1' and para.text.strip() == 'Table of Contents':
                in_toc = True
                continue
            if in_toc:
                # Stop at the next Heading 1 (which begins the body content)
                if para.style.name == 'Heading 1':
                    break
                # A TOC entry has visible text (non-empty after strip)
                if para.text.strip():
                    toc_entries.append(para.text.strip())

        toc_count = len(toc_entries)
        if toc_count >= 6:
            print(f"PASS: Component 2 — TOC has {toc_count} entries (expected >= 6) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — TOC has {toc_count} entries, expected 6. "
                  f"Entries found: {toc_entries}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: TOC includes an entry with 'Appendix A: Raw Data' text and a page number (0.2 points)
    #   The new TOC entry should reference the appendix heading with its page number.
    try:
        appendix_toc_entry = None
        in_toc = False
        for para in doc.paragraphs:
            if para.style.name == 'Heading 1' and para.text.strip() == 'Table of Contents':
                in_toc = True
                continue
            if in_toc:
                if para.style.name == 'Heading 1':
                    break
                if 'Appendix A' in para.text and 'Raw Data' in para.text:
                    appendix_toc_entry = para.text
                    break

        if appendix_toc_entry is not None:
            # Verify it also has a page number (contains digits after the tabs)
            parts = appendix_toc_entry.split('\t')
            has_page_number = any(part.strip().isdigit() for part in parts[1:])
            if has_page_number:
                page_num = [p.strip() for p in parts[1:] if p.strip().isdigit()]
                print(f"PASS: Component 3 — TOC entry 'Appendix A: Raw Data' found with "
                      f"page number {page_num} (0.2 pts)")
                total_score += 0.2
            else:
                print(f"PARTIAL: Component 3 — TOC entry 'Appendix A: Raw Data' found but "
                      f"no clear page number. Entry: {appendix_toc_entry!r}. "
                      f"Awarding 0.1 pts for entry presence only.")
                total_score += 0.1
        else:
            print("FAIL: Component 3 — No TOC entry found containing 'Appendix A: Raw Data'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
