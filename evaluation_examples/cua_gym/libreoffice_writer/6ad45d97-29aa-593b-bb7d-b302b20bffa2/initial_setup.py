"""
Initial Setup: Survey analysis document with TOC (5 entries) but NO bookmark at appendix heading.
Task ID: writer_struct_071
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_071'
OUTPUT = f'{WORKDIR}/survey_analysis.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc(doc, toc_title="Table of Contents"):
    """Insert a TOC field that LibreOffice can display/update."""
    # Add TOC heading
    toc_heading = doc.add_paragraph(toc_title, style='Heading 1')

    # Add the TOC field paragraph
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.space_before = Pt(0)
    toc_para.paragraph_format.space_after = Pt(0)

    run = toc_para.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)

    run2 = toc_para.add_run()
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-1" \\h \\z \\u '
    run2._r.append(instr_text)

    run3 = toc_para.add_run()
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fld_char_separate)

    # Static TOC entries (5 entries - no appendix)
    toc_entries = [
        ("Introduction", "2"),
        ("Survey Methodology", "3"),
        ("Key Findings", "4"),
        ("Demographic Analysis", "5"),
        ("Conclusions and Recommendations", "6"),
    ]

    for entry_text, page_num in toc_entries:
        entry_para = doc.add_paragraph(style='TOC 1')
        entry_para.clear()
        run_entry = entry_para.add_run(entry_text)
        run_entry.font.size = Pt(11)
        # Add tab and page number
        tab_run = entry_para.add_run(f'\t{page_num}')
        tab_run.font.size = Pt(11)

    run4 = toc_para.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run4._r.append(fld_char_end)


def add_section_content(doc, heading_text, content_paragraphs, level=1):
    """Add a heading and multiple paragraphs of content."""
    doc.add_heading(heading_text, level=level)
    for para_text in content_paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(8)


def create_initial():
    doc = Document()

    # Page setup: standard letter size
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title page
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run("Customer Satisfaction Survey Analysis")
    run.bold = True
    run.font.size = Pt(24)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle_para.add_run("Annual Report — Q4 2024")
    sub_run.font.size = Pt(14)

    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_before = Pt(36)
    auth_run = author_para.add_run("Research & Analytics Department\nOmega Solutions Inc.")
    auth_run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run("February 2025")
    date_run.font.size = Pt(12)

    doc.add_page_break()

    # TOC — lists exactly 5 heading entries (not appendix)
    toc_heading = doc.add_heading("Table of Contents", level=1)

    toc_entries = [
        ("Introduction", "2"),
        ("Survey Methodology", "3"),
        ("Key Findings", "4"),
        ("Demographic Analysis", "5"),
        ("Conclusions and Recommendations", "6"),
    ]
    for entry_text, page_num in toc_entries:
        p = doc.add_paragraph(style='Normal')
        run_e = p.add_run(f"{entry_text}")
        run_e.font.size = Pt(11)
        run_tab = p.add_run(f"\t\t\t{page_num}")
        run_tab.font.size = Pt(11)

    doc.add_page_break()

    # Section 1: Introduction (page 2)
    add_section_content(doc, "Introduction", [
        "This report presents the findings of our comprehensive customer satisfaction survey conducted "
        "during Q4 2024. A total of 2,847 respondents participated across five regional markets: "
        "North America, Europe, Asia-Pacific, Latin America, and the Middle East.",
        "The survey aimed to measure customer perceptions across four key dimensions: product quality, "
        "service responsiveness, value for money, and overall brand loyalty. Results indicate a "
        "significant improvement in customer satisfaction compared to the previous year's benchmark.",
        "This document is organized as follows: Section 2 describes our survey methodology, Section 3 "
        "presents key findings, Section 4 provides demographic breakdowns, and Section 5 offers "
        "strategic recommendations based on the data.",
    ])

    doc.add_page_break()

    # Section 2: Survey Methodology (page 3)
    add_section_content(doc, "Survey Methodology", [
        "The survey was administered via a mixed-method approach combining online questionnaires and "
        "telephone interviews. Participants were randomly selected from our active customer database, "
        "ensuring a stratified sample representative of our customer demographics.",
        "Data collection took place between October 1 and December 15, 2024. Survey instruments "
        "included 42 structured questions using Likert scales (1–5), multiple-choice items, and "
        "three open-ended response fields. The margin of error for reported statistics is ±2.1% "
        "at a 95% confidence level.",
        "All data was anonymized prior to analysis. Statistical analysis was performed using SPSS v28 "
        "and cross-validated with Python-based regression modeling. Qualitative responses were coded "
        "using thematic analysis by a team of four trained analysts.",
    ])

    doc.add_page_break()

    # Section 3: Key Findings (page 4)
    add_section_content(doc, "Key Findings", [
        "Overall customer satisfaction scores improved from 72.4% in Q4 2023 to 81.3% in Q4 2024, "
        "representing a statistically significant increase of 8.9 percentage points (p < 0.001).",
        "Product quality received the highest satisfaction ratings, with 87.2% of respondents rating "
        "it as 'Satisfied' or 'Very Satisfied'. Service responsiveness showed the greatest year-over-year "
        "improvement, rising from 65.1% to 78.4% — attributable to our enhanced 24/7 support rollout "
        "in March 2024.",
        "Net Promoter Score (NPS) increased from 34 to 52, placing us in the 'Good' category and "
        "approaching the industry average of 58 for enterprise software providers. Customer churn intent "
        "declined from 18.7% to 11.2% year-over-year.",
        "Key verbatim themes from open-ended responses included: (1) appreciation for faster issue "
        "resolution times, (2) requests for more intuitive UI in mobile applications, and (3) positive "
        "feedback on the newly launched self-service knowledge base.",
    ])

    doc.add_page_break()

    # Section 4: Demographic Analysis (page 5)
    add_section_content(doc, "Demographic Analysis", [
        "Respondents were segmented across four primary demographic categories: company size, industry "
        "vertical, geographic region, and account tenure. The largest segment comprised mid-market "
        "companies (50–500 employees), accounting for 44.3% of total responses.",
        "By industry, Technology & SaaS (28.1%), Financial Services (19.4%), Healthcare (15.7%), "
        "Retail & E-commerce (14.2%), and Manufacturing (12.6%) represented the top five verticals. "
        "Remaining 10% were distributed across education, government, and other sectors.",
        "Geographically, North America contributed 38.5% of responses, followed by Europe (29.3%), "
        "Asia-Pacific (21.8%), Latin America (6.9%), and Middle East & Africa (3.5%). Satisfaction "
        "scores were highest in Asia-Pacific (84.7%) and lowest in Latin America (74.2%), which "
        "aligns with known regional support capacity constraints.",
        "Long-tenure customers (5+ years) showed the highest satisfaction at 88.9%, while new customers "
        "(< 1 year) had the lowest at 73.4%. This gap suggests an opportunity to enhance onboarding "
        "programs and first-year customer success touchpoints.",
    ])

    doc.add_page_break()

    # Section 5: Conclusions (page 6)
    add_section_content(doc, "Conclusions and Recommendations", [
        "The data clearly demonstrates meaningful progress in customer satisfaction across all measured "
        "dimensions. The 8.9-point overall increase reflects the cumulative impact of strategic "
        "investments in support infrastructure, product improvements, and account management training.",
        "Based on our analysis, we recommend the following priority initiatives for 2025: First, launch "
        "a structured new-customer onboarding program targeting the 73.4% satisfaction gap among "
        "first-year accounts. Second, accelerate mobile UX improvements addressing the top verbatim "
        "complaint theme. Third, expand the APAC support team to sustain the region's above-average "
        "satisfaction trajectory.",
        "With disciplined execution of these recommendations, we project overall satisfaction will reach "
        "87% by Q4 2025, with NPS climbing to 60–65, consistent with industry top-quartile performance.",
    ])

    doc.add_page_break()

    # Appendix A: Raw Data (page 7) — Heading 1, NO bookmark
    doc.add_heading("Appendix A: Raw Data", level=1)

    appendix_intro = doc.add_paragraph(
        "This appendix contains the complete raw data tables referenced throughout the report. "
        "Data is presented in tabular form by survey dimension and demographic segment."
    )

    # Add a table with some raw data
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ["Region", "Respondents", "Avg Score", "NPS"]
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    data_rows = [
        ["North America", "1,096", "82.1", "54"],
        ["Europe", "834", "81.8", "51"],
        ["Asia-Pacific", "621", "84.7", "58"],
        ["Latin America", "196", "74.2", "43"],
        ["Middle East & Africa", "100", "78.9", "47"],
    ]
    for row_idx, row_data in enumerate(data_rows, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph("")

    # More appendix content to fill page 7
    add_section_content(doc, "A.1 Survey Instrument Summary", [
        "The full 42-question survey instrument is available upon request from the Research & "
        "Analytics Department. Question categories included: Product Quality (Q1–Q10), Service "
        "Experience (Q11–Q20), Value Perception (Q21–Q28), Brand Loyalty (Q29–Q36), and "
        "Demographic Information (Q37–Q42).",
    ], level=2)

    doc.add_page_break()

    # Page 8 — more appendix data
    add_section_content(doc, "A.2 Verbatim Response Themes", [
        "The following table summarizes the most frequently coded themes from open-ended survey "
        "responses. Thematic coding was performed by four independent analysts with inter-rater "
        "reliability (Cohen's κ) of 0.82.",
    ], level=2)

    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    headers2 = ["Theme", "Frequency", "Sentiment"]
    for col_idx, header in enumerate(headers2):
        cell = table2.cell(0, col_idx)
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    themes = [
        ["Faster support response times", "487", "Positive"],
        ["Mobile app usability issues", "312", "Negative"],
        ["Self-service knowledge base praise", "298", "Positive"],
        ["Pricing transparency requests", "241", "Neutral"],
        ["Feature request: API improvements", "189", "Neutral"],
    ]
    for row_idx, row_data in enumerate(themes, 1):
        for col_idx, value in enumerate(row_data):
            cell = table2.cell(row_idx, col_idx)
            cell.text = value
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
