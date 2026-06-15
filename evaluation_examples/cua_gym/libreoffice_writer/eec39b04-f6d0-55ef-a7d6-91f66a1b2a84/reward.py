"""
FINAL REWARD SCRIPT - SUCCESS
Task: Could you include "Hooper, L., et al. (2012). Effects of chocolate on blood pressure and cardiovascular risk. Cochrane Database of Systematic Reviews, (8), CD008893." in my reference list, then insert the reference number as a superscript where I marked "<cite>" in the conclusion’s first paragraph?
Generated: 2025-10-14 12:14:24
Status: success
Model: azure-o3
Total Steps: 3
"""

import os
from docx import Document

"""
Reward Script for Writer Task Verification
Task: The user asked to add a new reference (Hooper et al., 2012) to the reference list and to insert the corresponding reference number as a superscript in the first paragraph of the Conclusion section (where <cite> was previously marked).

The script awards points ONLY for the two concrete objectives:
1. The Hooper (2012) reference is present in the reference list (0.4 pts)
2. A superscript number is present in the first Conclusion paragraph (0.3 pts)
3. The superscript number matches the numerical position of the Hooper reference in the list (0.3 pts)

Total possible = 1.0
The script prints detailed diagnostics and finally prints "REWARD: X.X".
"""

def verify_task(file_path: str) -> float:
    score = 0.0

    # --- Load document ------------------------------------------------------
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Could not load document: {e}")
        return 0.0

    # --- Locate reference list ---------------------------------------------
    references_start = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == "references":
            references_start = idx
            break

    hooper_index = None  # 1-based position of Hooper ref in list
    if references_start is not None:
        print(f"✓ References heading found at paragraph {references_start}")
        # Collect consecutive non-empty paragraphs after "References" until next heading
        references = []
        for p in doc.paragraphs[references_start + 1:]:
            # Stop if another heading encountered
            if p.style.name.lower().startswith("heading") and p.text.strip():
                break
            txt = p.text.strip()
            if txt:
                references.append(txt)
        print(f"Reference list size: {len(references)}")
        # Find Hooper 2012 reference in list
        for i, ref in enumerate(references):
            if "hooper" in ref.lower() and "2012" in ref:
                hooper_index = i + 1  # convert to 1-based index
                break
        if hooper_index is not None:
            print(f"✓ Hooper reference found at position {hooper_index}")
            score += 0.4
        else:
            print("✗ Hooper 2012 reference not found in list")
    else:
        print("✗ References section not found")

    # --- Locate Conclusion first paragraph ---------------------------------
    conclusion_idx = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == "conclusion" and p.style.name.lower().startswith("heading"):
            conclusion_idx = idx
            break

    if conclusion_idx is None:
        print("✗ Conclusion heading not found")
    else:
        # Find first non-empty paragraph after the heading
        para_idx = conclusion_idx + 1
        while para_idx < len(doc.paragraphs) and not doc.paragraphs[para_idx].text.strip():
            para_idx += 1
        if para_idx >= len(doc.paragraphs):
            print("✗ No paragraph found after Conclusion heading")
        else:
            para = doc.paragraphs[para_idx]
            # Collect superscript runs (Writer exports them with run.font.superscript == True)
            superscripts = [run.text.strip() for run in para.runs if run.font.superscript and run.text.strip()]
            if superscripts:
                print(f"✓ Superscript citations in Conclusion paragraph: {superscripts}")
                # Extract numeric part of first superscript
                num_part = ''.join(ch for ch in superscripts[0] if ch.isdigit())
                if num_part.isdigit():
                    citation_num = int(num_part)
                    print(f"First citation number: {citation_num}")
                    score += 0.3  # superscript exists
                    # Check if number matches Hooper position
                    if hooper_index is not None and citation_num == hooper_index:
                        print("✓ Superscript number matches Hooper reference index")
                        score += 0.3
                    else:
                        if hooper_index is not None:
                            print("✗ Superscript number does not match Hooper reference index")
                else:
                    print("✗ Superscript does not contain a numeric citation")
            else:
                print("✗ No superscript citation found in Conclusion paragraph")

    # --- Final score --------------------------------------------------------
    final_score = min(score, 1.0)
    print(f"Final score: {final_score}")
    return final_score


if __name__ == "__main__":
    FILE_PATH = "/home/user/could_you_include_hooper_l_et_al_2012_effects_of_chocolate_on_blood_pressure_and_cardiovascular_risk.docx"
    reward = verify_task(FILE_PATH)
    print(f"REWARD: {reward}")
