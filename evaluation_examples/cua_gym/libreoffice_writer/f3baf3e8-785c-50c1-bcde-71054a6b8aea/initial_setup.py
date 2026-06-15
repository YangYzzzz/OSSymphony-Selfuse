"""
Initial Setup: Math textbook chapter 5 document with bookmark but no cross-reference
Task ID: writer_struct_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_024'
OUTPUT = f'{WORKDIR}/Desktop/math_textbook_ch5.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bookmark(paragraph, bookmark_name):
    """Insert a bookmark around the entire paragraph content."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")

    # Create bookmark start element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '1')
    bookmark_start.set(qn('w:name'), bookmark_name)

    # Create bookmark end element
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '1')

    # Insert at the beginning of the paragraph
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)


def add_heading(doc, text, level):
    """Add a heading with given text and level."""
    p = doc.add_heading(text, level=level)
    return p


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def create_initial():
    """Create the initial math textbook chapter 5 document."""
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Set default font and margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---- PAGE 1 ----
    # Chapter title
    title = doc.add_heading('Chapter 5: The Golden Ratio and Its Applications', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    intro = doc.add_paragraph(
        'This chapter explores one of the most fascinating mathematical constants known to humanity: '
        'the Golden Ratio. Represented by the Greek letter φ (phi), this irrational number approximately '
        'equal to 1.618 has appeared in art, architecture, nature, and mathematics for thousands of years.'
    )

    doc.add_paragraph(
        'We will examine the mathematical properties of φ, its relationship to the Fibonacci sequence, '
        'and its remarkable appearances in natural phenomena and human artistic creation. '
        'By the end of this chapter, students will have a deep understanding of why this ratio is '
        'considered the most aesthetically pleasing proportion in the visual arts.'
    )

    # Section 5.1
    doc.add_heading('5.1 Introduction to Mathematical Proportions', level=1)

    doc.add_paragraph(
        'Before we can appreciate the Golden Ratio, we must understand the concept of mathematical '
        'proportions. A proportion is a relationship between two quantities, often expressed as a fraction '
        'or ratio. In ancient Greece, mathematicians were particularly interested in proportions that '
        'appeared both in geometry and in the natural world.'
    )

    doc.add_paragraph(
        'The study of proportions dates back to Euclid of Alexandria (circa 300 BCE), whose monumental '
        'work "Elements" laid the foundation for much of classical geometry. Euclid was particularly '
        'interested in what he called "extreme and mean ratio" — a division of a line segment such that '
        'the ratio of the whole segment to the larger part equals the ratio of the larger part to the '
        'smaller part.'
    )

    doc.add_paragraph(
        'This seemingly simple geometric relationship would go on to fascinate mathematicians, artists, '
        'and architects for the next two millennia. The concept was later named the "Golden Ratio" by '
        'Renaissance mathematicians, who believed it represented a divinely inspired proportion found '
        'throughout creation.'
    )

    # Add page break to go to page 2
    doc.add_page_break()

    # ---- PAGE 2 ----
    doc.add_heading('5.2 Algebraic Definition of the Golden Ratio', level=1)

    doc.add_paragraph(
        'The Golden Ratio can be defined algebraically through a simple but profound equation. '
        'Consider a line segment of total length (a + b), divided into two parts a and b, '
        'where a > b. We say the line is divided in the Golden Ratio if:'
    )

    eq_para = doc.add_paragraph()
    eq_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    eq_run = eq_para.add_run('(a + b) / a = a / b = φ')
    eq_run.bold = True
    eq_run.font.size = Pt(14)

    doc.add_paragraph(
        'Solving this equation algebraically, we find that φ satisfies the quadratic equation:'
    )

    eq_para2 = doc.add_paragraph()
    eq_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    eq_run2 = eq_para2.add_run('φ² = φ + 1')
    eq_run2.bold = True
    eq_run2.font.size = Pt(14)

    doc.add_paragraph(
        'Applying the quadratic formula, we obtain the exact value of the Golden Ratio:'
    )

    eq_para3 = doc.add_paragraph()
    eq_para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    eq_run3 = eq_para3.add_run('φ = (1 + √5) / 2 ≈ 1.6180339887...')
    eq_run3.bold = True
    eq_run3.font.size = Pt(14)

    doc.add_paragraph(
        'This value is irrational, meaning it cannot be expressed as a simple fraction. '
        'Its decimal expansion never terminates and never repeats. '
        'Interestingly, φ is sometimes called the "most irrational" number because '
        'it is the hardest to approximate using simple fractions — a property with deep connections '
        'to continued fractions and Diophantine approximation theory.'
    )

    doc.add_paragraph(
        'The reciprocal of φ is also remarkable: 1/φ = φ - 1 ≈ 0.6180339887... '
        'This means the decimal parts of φ and 1/φ are identical, a unique property '
        'not shared by any other positive real number.'
    )

    doc.add_paragraph(
        'Furthermore, φ squared equals φ plus 1: φ² = φ + 1 ≈ 2.6180339887... '
        'These self-referential properties make the Golden Ratio particularly special '
        'in the field of recursive mathematical structures.'
    )

    doc.add_page_break()

    # ---- PAGE 3 ----
    doc.add_heading('5.3 The Golden Ratio and Fibonacci Sequence', level=1)

    doc.add_paragraph(
        'Perhaps the most remarkable property of the Golden Ratio is its intimate connection '
        'to the Fibonacci sequence. The Fibonacci sequence, named after the Italian mathematician '
        'Leonardo of Pisa (known as Fibonacci), is defined by the recurrence relation:'
    )

    eq_para4 = doc.add_paragraph()
    eq_para4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    eq_run4 = eq_para4.add_run('F(n) = F(n-1) + F(n-2), with F(0) = 0, F(1) = 1')
    eq_run4.bold = True

    doc.add_paragraph(
        'The first several Fibonacci numbers are: 0, 1, 1, 2, 3, 5, 8, 13, 21, 34, 55, 89, 144, ...'
    )

    # Add a table showing Fibonacci ratios
    doc.add_paragraph('As we compute successive ratios F(n+1)/F(n), we observe convergence to φ:')

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    headers = ['n', 'F(n)', 'F(n+1)/F(n)']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    fib_data = [
        ('5', '5', '1.6000000'),
        ('6', '8', '1.6250000'),
        ('7', '13', '1.6153846'),
        ('8', '21', '1.6190476'),
        ('9', '34', '1.6176471'),
        ('10', '55', '1.6181818'),
        ('20', '6765', '1.6180340'),
    ]
    for i, (n, fn, ratio) in enumerate(fib_data, 1):
        table.cell(i, 0).text = n
        table.cell(i, 1).text = fn
        table.cell(i, 2).text = ratio

    doc.add_paragraph()
    doc.add_paragraph(
        'This convergence is not a coincidence. It can be proven rigorously using the Binet formula, '
        'which gives an explicit expression for the nth Fibonacci number in terms of φ:'
    )

    eq_para5 = doc.add_paragraph()
    eq_para5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    eq_run5 = eq_para5.add_run('F(n) = (φⁿ - ψⁿ) / √5, where ψ = (1 - √5) / 2 ≈ -0.618')
    eq_run5.bold = True

    doc.add_page_break()

    # ---- PAGE 4 ----
    # This page contains the bookmark 'golden_ratio_definition'
    doc.add_heading('5.4 Formal Definition and Key Properties', level=1)

    # The bookmark paragraph - the formal definition
    defn_para = doc.add_paragraph(
        'The Golden Ratio, denoted φ (phi), is defined as the unique positive real number satisfying '
        'the equation φ² = φ + 1. Its exact value is φ = (1 + √5) / 2 ≈ 1.6180339887. '
        'This number possesses the extraordinary property that its reciprocal differs from it only by 1: '
        '1/φ = φ - 1. No other positive number shares this self-referential characteristic.'
    )

    # Add bookmark to this paragraph
    add_bookmark(defn_para, 'golden_ratio_definition')

    doc.add_paragraph(
        'The key algebraic properties of the Golden Ratio include:'
    )

    # Properties as a bulleted list
    props = [
        'φ² = φ + 1 (fundamental defining property)',
        '1/φ = φ - 1 ≈ 0.6180339887',
        'φ = 1 + 1/φ (continued fraction representation)',
        'φⁿ = φⁿ⁻¹ + φⁿ⁻² (power recurrence)',
        'φ is an algebraic integer of degree 2',
        'The continued fraction expansion of φ is [1; 1, 1, 1, 1, ...] — all ones',
    ]
    for prop in props:
        doc.add_paragraph(prop, style='List Bullet')

    doc.add_paragraph(
        'These properties make φ the foundation of numerous mathematical constructions, '
        'from Penrose tilings to quasi-crystals and beyond. The formal definition above '
        'distinguishes φ from its conjugate ψ = (1 - √5)/2 ≈ -0.618, which also satisfies '
        'ψ² = ψ + 1 but is negative.'
    )

    doc.add_paragraph(
        'In the next sections, we will explore how this definition manifests in geometric '
        'constructions, particularly in the golden rectangle and the golden spiral, '
        'which have been used in artistic composition for centuries.'
    )

    doc.add_page_break()

    # ---- PAGE 5 ----
    doc.add_heading('5.5 The Golden Rectangle', level=1)

    doc.add_paragraph(
        'A golden rectangle is a rectangle whose side lengths are in the Golden Ratio. '
        'That is, if the longer side has length a and the shorter side has length b, '
        'then a/b = φ. Golden rectangles have a remarkable self-similar property: '
        'if you remove a square from one end of a golden rectangle, the remaining rectangle '
        'is also a golden rectangle.'
    )

    doc.add_paragraph(
        'This property can be proven easily. If we have a golden rectangle with sides φ × 1, '
        'and we remove a 1 × 1 square, we are left with a rectangle of sides 1 × (φ - 1). '
        'Since φ - 1 = 1/φ, this smaller rectangle has sides 1 × (1/φ), or equivalently φ × 1 '
        'after rotating — exactly another golden rectangle.'
    )

    doc.add_paragraph(
        'The process can be repeated indefinitely, generating an infinite sequence of '
        'ever-smaller golden rectangles. Connecting the corners of these squares with a '
        'smooth curve produces the famous golden spiral, an approximation of the logarithmic '
        'spiral with growth factor φ.'
    )

    doc.add_paragraph(
        'Construction of a golden rectangle using compass and straightedge:'
    )

    steps = [
        'Begin with a unit square ABCD.',
        'Find the midpoint M of side AB.',
        'Draw an arc centered at M with radius MC, intersecting the extension of AB at point E.',
        'Erect a perpendicular to AE at E, meeting the extension of DC at F.',
        'Rectangle AEFD is a golden rectangle with AE/AD = φ.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    doc.add_paragraph(
        'This classical construction demonstrates that the Golden Ratio is a constructible number, '
        'meaning it can be produced using only a compass and straightedge starting from a unit length.'
    )

    doc.add_page_break()

    # ---- PAGE 6 ----
    doc.add_heading('5.6 The Golden Ratio in Architecture', level=1)

    doc.add_paragraph(
        'The Golden Ratio has been claimed to appear in numerous architectural masterpieces '
        'throughout history. While some of these claims are contested by modern scholars, '
        'there is substantial evidence that Renaissance architects deliberately incorporated '
        'φ into their designs.'
    )

    doc.add_paragraph(
        'The Parthenon in Athens, built between 447 and 432 BCE, is perhaps the most famous '
        'example. The ratio of its width to height, as well as numerous internal proportions, '
        'closely approximate the Golden Ratio. Whether this was intentional or the result of '
        'general aesthetic principles remains a subject of scholarly debate.'
    )

    doc.add_paragraph(
        'Leonardo da Vinci\'s collaborations with Luca Pacioli — who wrote "De Divina Proportione" '
        'in 1509 — represent perhaps the most documented use of φ in Renaissance art and architecture. '
        'Pacioli explicitly advocated for the use of the Golden Ratio in artistic composition, '
        'and da Vinci provided the illustrations for his treatise.'
    )

    doc.add_paragraph(
        'In modern architecture, the Swiss-French architect Le Corbusier developed the "Modulor" '
        'system, a scale of proportions based on human measurements and the Golden Ratio. '
        'This system was used in the design of the Unité d\'Habitation in Marseille (1952) '
        'and the Chandigarh Capitol Complex in India.'
    )

    doc.add_paragraph(
        'The United Nations Secretariat Building in New York City (1952) also incorporates '
        'the Golden Ratio in its proportions, a tribute to Le Corbusier\'s influence on '
        'mid-century modernist architecture.'
    )

    doc.add_page_break()

    # ---- PAGE 7 ----
    doc.add_heading('5.7 The Golden Ratio in Nature', level=1)

    doc.add_paragraph(
        'Perhaps the most compelling evidence for the universality of the Golden Ratio comes '
        'from its appearances in the natural world. Botanists have long observed that the '
        'arrangement of leaves around plant stems (phyllotaxis) frequently follows Fibonacci '
        'numbers, which in turn converge to the Golden Ratio.'
    )

    doc.add_paragraph(
        'Sunflower seed heads provide a striking example. The seeds are arranged in two sets '
        'of interlocking spirals: one set curving clockwise, the other counterclockwise. '
        'The number of spirals in each direction are typically consecutive Fibonacci numbers, '
        'such as 34 and 55, or 55 and 89. This arrangement allows for maximum packing density.'
    )

    doc.add_paragraph(
        'The nautilus shell is another classic example. Its growth follows a logarithmic spiral '
        'with a growth factor approximately equal to φ per 90 degrees of rotation. '
        'While the precise ratio varies among individual shells and species, '
        'the overall form approximates a golden spiral.'
    )

    doc.add_paragraph(
        'DNA molecules, at the molecular scale, have been noted to exhibit proportions related to '
        'the Golden Ratio. The width of the DNA double helix is approximately 34 Ångströms, '
        'while one full twist is approximately 21 Ångströms long — both Fibonacci numbers, '
        'with a ratio of 34/21 ≈ 1.619.'
    )

    doc.add_paragraph(
        'In the animal kingdom, the proportions of the human body have been extensively analyzed '
        'for Golden Ratio relationships. While many popular claims are exaggerated or based on '
        'selective measurement, statistical studies do suggest that certain facial proportions '
        'considered most attractive cluster around the Golden Ratio.'
    )

    doc.add_page_break()

    # ---- PAGE 8 ----
    doc.add_heading('5.8 The Golden Ratio in Art and Music', level=1)

    doc.add_paragraph(
        'The influence of the Golden Ratio in visual art is profound and well-documented. '
        'Artists from the Renaissance through the modern era have consciously used φ as a '
        'compositional tool, believing it produces works that are inherently more beautiful '
        'and harmonious to the human eye.'
    )

    doc.add_paragraph(
        'Sandro Botticelli\'s "The Birth of Venus" (circa 1486) exhibits numerous Golden Ratio '
        'proportions in the figure of Venus herself. The ratio of her height to the position '
        'of her navel, and the proportions of her face, closely approximate φ. '
        'Botticelli was a contemporary of Pacioli and was likely familiar with the mathematical '
        'treatises of his day.'
    )

    doc.add_paragraph(
        'Salvador Dalí explicitly used the Golden Ratio in "The Sacrament of the Last Supper" (1955). '
        'The dimensions of the canvas itself are in the golden ratio, and the large dodecahedron '
        'visible in the upper portion of the painting (a solid associated with φ) '
        'reinforces this mathematical theme.'
    )

    doc.add_paragraph(
        'In music, the composer Béla Bartók used the Fibonacci sequence and Golden Ratio '
        'extensively in his compositions. Analysis of his "Music for Strings, Percussion and Celesta" '
        'reveals that the climax of the piece occurs at approximately the golden section '
        'of its total duration — a technique he used repeatedly in his work.'
    )

    doc.add_paragraph(
        'The composer Debussy is also said to have used the Golden Ratio in works such as '
        '"La Mer" and various Preludes, although scholarly analysis of this claim has produced '
        'mixed results. The relationship between mathematical proportion and perceived musical '
        'beauty remains an active area of musicological research.'
    )

    doc.add_page_break()

    # ---- PAGE 9 ----
    doc.add_heading('5.9 Cross-References and Revisiting Key Concepts', level=1)

    doc.add_paragraph(
        'As we progress through the applications of the Golden Ratio, it becomes important to '
        'revisit some of the foundational definitions established earlier in this chapter. '
        'The formal mathematical properties of φ underpin all of the geometric and aesthetic '
        'applications we have examined.'
    )

    doc.add_paragraph(
        'The self-similar property of the golden rectangle, for instance, derives directly '
        'from the algebraic identity 1/φ = φ - 1. Without the precise value and defining '
        'equation of φ, we could not construct golden rectangles with compass and straightedge, '
        'nor could we prove that the removed square always leaves another golden rectangle.'
    )

    # The critical paragraph - sentence ends at "Recall the definition from page "
    # NO cross-reference here in the initial state
    recall_para = doc.add_paragraph(
        'The nautilus spiral approximation and the phyllotaxis patterns in sunflowers both '
        'depend on the mathematical properties of φ established in the early sections of this chapter. '
        'Recall the definition from page '
    )
    # NOTE: No cross-reference field here — that is what the task asks the agent to add

    doc.add_paragraph(
        'The continued fraction representation [1; 1, 1, 1, ...] explains why Fibonacci numbers '
        'provide the best rational approximations to φ, which in turn explains why plants using '
        'Fibonacci phyllotaxis achieve optimal packing — each new leaf or seed is positioned at '
        'approximately 137.5 degrees (the golden angle, 360°/φ²) from the previous one.'
    )

    doc.add_paragraph(
        'Understanding this connection between the abstract algebraic definition and its '
        'physical manifestations is central to appreciating the depth of the Golden Ratio\'s '
        'significance in mathematics and the natural sciences.'
    )

    doc.add_page_break()

    # ---- PAGE 10 ----
    doc.add_heading('5.10 Quasi-Crystals and Penrose Tilings', level=1)

    doc.add_paragraph(
        'One of the most surprising modern applications of the Golden Ratio is in the theory of '
        'quasi-crystals and Penrose tilings. In 1974, the mathematician and physicist Roger Penrose '
        'discovered a way to tile the plane using only two shapes — "kite" and "dart" — that never '
        'repeats periodically but nonetheless fills the entire plane without gaps or overlaps.'
    )

    doc.add_paragraph(
        'The Penrose tiling has remarkable properties: it exhibits five-fold rotational symmetry '
        '(impossible for ordinary periodic crystals) and the ratio of the two tile types in any '
        'sufficiently large region is always φ. The shape of the tiles themselves is derived from '
        'the geometry of the regular pentagon, which is intimately related to φ.'
    )

    doc.add_paragraph(
        'In 1984, the materials scientist Dan Shechtman discovered a real metallic alloy '
        '(aluminum-manganese) whose X-ray diffraction pattern showed five-fold symmetry — '
        'a result considered impossible by classical crystallography. This discovery, for which '
        'Shechtman received the 2011 Nobel Prize in Chemistry, opened the field of quasi-crystals.'
    )

    doc.add_paragraph(
        'The internal structure of quasi-crystals like icosahedral Al-Mn is governed by the '
        'Golden Ratio: atomic positions along any axis repeat at intervals in the ratio φ. '
        'This aperiodic but ordered structure is a three-dimensional analog of Penrose tiling.'
    )

    doc.add_paragraph(
        'The mathematical framework for understanding quasi-crystals draws heavily on '
        'the theory of cut-and-project tilings, which uses higher-dimensional lattices '
        'projected onto lower-dimensional spaces. The Golden Ratio appears naturally '
        'because it is the eigenvalue of the "inflation" symmetry of Penrose tilings.'
    )

    doc.add_page_break()

    # ---- PAGE 11 ----
    doc.add_heading('5.11 The Golden Ratio in Modern Mathematics', level=1)

    doc.add_paragraph(
        'Beyond its classical applications, the Golden Ratio continues to appear in cutting-edge '
        'mathematical research. In number theory, φ plays a central role in the theory of '
        'continued fractions and Diophantine approximation.'
    )

    doc.add_paragraph(
        'Hurwitz\'s theorem states that for any irrational number α, there are infinitely many '
        'rational fractions p/q satisfying |α - p/q| < 1/(√5 · q²). Furthermore, the constant √5 '
        'in this bound is optimal — it cannot be replaced by any larger constant if we require '
        'the inequality to hold for infinitely many p/q. The extremal case is achieved precisely '
        'by α = φ, confirming φ\'s status as the "most irrational" number.'
    )

    doc.add_paragraph(
        'In dynamical systems, the Golden Ratio appears in KAM theory (Kolmogorov-Arnold-Moser), '
        'which studies the stability of nearly-integrable Hamiltonian systems. '
        'Invariant tori with rotation number φ are the most robust against perturbation — '
        'again because φ is hardest to approximate by rationals.'
    )

    doc.add_paragraph(
        'In graph theory, the Golden Ratio appears as the spectral gap of certain Ramanujan graphs '
        'and in the theory of expander graphs. The Petersen graph and related structures have '
        'eigenvalues involving φ.'
    )

    doc.add_paragraph(
        'Algebraic geometry connects φ to the regular icosahedron and dodecahedron — '
        'the two Platonic solids with five-fold symmetry. The coordinates of their vertices '
        'can all be expressed in terms of φ, and their symmetry group (the icosahedral group A₅) '
        'is deeply connected to the algebraic properties of φ.'
    )

    doc.add_page_break()

    # ---- PAGE 12 ----
    doc.add_heading('5.12 Summary and Exercises', level=1)

    doc.add_paragraph(
        'In this chapter, we have explored the Golden Ratio from multiple perspectives: '
        'algebraic, geometric, historical, artistic, and scientific. The Golden Ratio φ ≈ 1.618 '
        'emerges from a simple geometric problem posed by the ancient Greeks, yet its implications '
        'reach into virtually every field of human inquiry.'
    )

    doc.add_paragraph(
        'Key concepts covered in this chapter include:'
    )

    summary_items = [
        'The algebraic definition of φ and its key properties (φ² = φ + 1, 1/φ = φ - 1)',
        'The connection between φ and the Fibonacci sequence via the Binet formula',
        'Geometric constructions: the golden rectangle and golden spiral',
        'Historical and artistic applications from ancient Greece to modern architecture',
        'Natural occurrences of φ in plant growth (phyllotaxis), shells, and DNA',
        'Advanced mathematical appearances in quasi-crystals, Penrose tilings, and number theory',
    ]
    for item in summary_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Exercises', level=2)

    exercises = [
        'Prove that the continued fraction [1; 1, 1, 1, ...] converges to φ. '
        '(Hint: let x = 1 + 1/x and solve the resulting quadratic equation.)',
        'Use the Binet formula to show that F(n) is the nearest integer to φⁿ/√5 for all n ≥ 1.',
        'Construct a golden rectangle using compass and straightedge, following the steps '
        'outlined in Section 5.5. Verify that the ratio of the sides equals φ to three decimal places.',
        'Count the spiral arms in a photograph of a sunflower head. Verify that the numbers '
        'in the two directions are consecutive Fibonacci numbers.',
        'Research one additional appearance of the Golden Ratio in architecture or art not '
        'mentioned in this chapter. Write a one-paragraph analysis of how φ appears.',
        'Using the algebraic identity φ² = φ + 1, derive the continued fraction expansion '
        'φ = 1 + 1/(1 + 1/(1 + 1/...)) without using the limit definition.',
    ]
    for i, ex in enumerate(exercises, 1):
        doc.add_paragraph(f'{i}. {ex}')
        doc.add_paragraph()

    doc.add_paragraph(
        '— End of Chapter 5 —'
    ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
