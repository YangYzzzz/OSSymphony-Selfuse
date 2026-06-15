"""
Initial Setup: Create mail merge name badges at a conference
Task ID: writer_mt_025
Domain: libreoffice_writer

Creates:
  1. A CSV data source with 120 conference registrant records
  2. A blank Writer document (the starting point for the mail merge task)
  3. Opens the document in LibreOffice Writer
"""

import csv
import os
import random
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_025'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
CSV_PATH = f'{WORKDIR}/ConferenceRegistrants.csv'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# --- Realistic name pools ---
FIRST_NAMES = [
    "Sarah", "Marcus", "Elena", "James", "Priya", "David", "Maria", "Chen",
    "Aisha", "Robert", "Yuki", "Thomas", "Fatima", "Michael", "Sofia", "Wei",
    "Amara", "Daniel", "Ingrid", "Carlos", "Naomi", "Patrick", "Leila", "Anton",
    "Zara", "Benjamin", "Mei", "Oliver", "Sana", "Lucas", "Anya", "Raj",
    "Hannah", "Felix", "Diana", "Ahmed", "Julia", "Leo", "Nadia", "Victor",
    "Chloe", "Ivan", "Rosa", "Nathan", "Elif", "George", "Mira", "Samuel",
    "Tanya", "Kevin", "Lina", "Oscar", "Freya", "Hassan", "Emma", "Rafael",
    "Vera", "Simon", "Kira", "Adam"
]

LAST_NAMES = [
    "Chen", "Johnson", "Patel", "Williams", "Nakamura", "Rodriguez", "Kim",
    "O'Brien", "Okafor", "Larsson", "Singh", "Martinez", "Tanaka", "Brown",
    "Muller", "da Silva", "Park", "Thompson", "Al-Rashid", "Johansson",
    "Gupta", "Davis", "Fernandez", "Watanabe", "Clark", "Novak", "Sharma",
    "Wilson", "Petrov", "Garcia", "Yamamoto", "Taylor", "Kowalski", "Lee",
    "Andersen", "Dubois", "Reyes", "Fischer", "Santos", "Cooper",
    "Bergstrom", "Li", "Moore", "Volkov", "Rossi", "Chang", "Martin",
    "Svensson", "Abbas", "Nguyen", "Schmidt", "Jackson", "Katz", "Aoki",
    "Rivera", "Hansen", "Ito", "Foster", "Berg", "Torres"
]

ORGANIZATIONS = [
    "MIT Lincoln Laboratory", "Google DeepMind", "Stanford NLP Group",
    "Meta AI Research", "Carnegie Mellon University", "OpenAI",
    "Microsoft Research", "UC Berkeley EECS", "Amazon Science",
    "NVIDIA Research", "IBM Watson", "Princeton University",
    "ETH Zurich", "Max Planck Institute", "University of Oxford",
    "Cambridge AI Lab", "KAIST", "University of Tokyo",
    "Tsinghua University", "Peking University", "INRIA",
    "Allen Institute for AI", "Anthropic", "Cohere",
    "Salesforce Research", "Adobe Research", "Apple ML Research",
    "Samsung AI Center", "Huawei Noah's Ark Lab", "Baidu Research",
    "Toyota Research Institute", "Bosch Center for AI",
    "Siemens Technology", "Intel Labs", "Qualcomm AI Research",
    "Johns Hopkins University", "University of Washington",
    "Georgia Tech", "University of Michigan", "NYU Courant Institute",
    "Columbia University", "University of Toronto", "MILA Quebec",
    "University of Montreal", "Technical University of Munich",
    "Imperial College London", "University of Edinburgh",
    "National University of Singapore", "HKUST", "Seoul National University"
]

REG_TYPES = ["Speaker", "Attendee", "Workshop Leader", "Poster Presenter", "Sponsor", "VIP"]


def generate_registrants(n=120):
    """Generate n realistic conference registrant records."""
    random.seed(42)
    records = []
    used_names = set()
    for _ in range(n):
        while True:
            first = random.choice(FIRST_NAMES)
            last = random.choice(LAST_NAMES)
            full = f"{first} {last}"
            if full not in used_names:
                used_names.add(full)
                break
        org = random.choice(ORGANIZATIONS)
        reg_type = random.choice(REG_TYPES)
        email = f"{first.lower()}.{last.lower().replace(' ', '')}@{org.lower().replace(' ', '').replace(',', '')[:20]}.org"
        records.append({
            "AttendeeName": full,
            "Organization": org,
            "RegistrationType": reg_type,
            "Email": email
        })
    return records


def create_csv(records):
    """Write registrant records to CSV data source."""
    with open(CSV_PATH, 'w', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=["AttendeeName", "Organization", "RegistrationType", "Email"])
        writer.writeheader()
        writer.writerows(records)
    print(f"CSV data source created: {CSV_PATH} ({len(records)} records)")


def create_initial_doc():
    """Create a minimal Writer document as the starting point for mail merge."""
    doc = Document()

    # Set up standard page margins
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Add a simple title/instruction paragraph (this is what the agent sees initially)
    title = doc.add_heading("Conference Name Badges", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    note = doc.add_paragraph()
    note.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = note.add_run("Data source: ConferenceRegistrants.csv (120 records)")
    run.font.size = Pt(11)
    run.font.italic = True

    note2 = doc.add_paragraph()
    note2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run2 = note2.add_run("Fields available: AttendeeName, Organization, RegistrationType, Email")
    run2.font.size = Pt(11)
    run2.font.italic = True

    doc.save(OUTPUT)
    print(f"Initial document created: {OUTPUT}")


def main():
    records = generate_registrants(120)
    create_csv(records)
    create_initial_doc()

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


main()
