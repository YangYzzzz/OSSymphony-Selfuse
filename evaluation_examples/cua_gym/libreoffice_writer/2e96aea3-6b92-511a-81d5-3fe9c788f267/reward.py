"""
Reward Script: Mail merge conference name badges with 2-column layout
Task ID: writer_mt_025
Domain: libreoffice_writer
Scoring:
  Component 1: 2-column page layout (0.20 pts)
  Component 2: 120 attendee names in ~24pt bold centered text (0.30 pts)
  Component 3: 120 organization names in ~16pt centered text (0.20 pts)
  Component 4: Page/column breaks for badge separation (0.15 pts)
  Component 5: Correct badge structure — name+org pairs (0.15 pts)
"""

import os

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_025'


def verify_task(file_path):
    """
    Verify mail merge name badge document with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 2-column page layout (0.20 points)
    # The golden doc has w:cols with w:num="2". Initial has no num attribute.
    try:
        sec = doc.sections[0]
        cols_elem = sec._sectPr.find(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols'
        )
        col_num = None
        if cols_elem is not None:
            num_attr = cols_elem.attrib.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'
            )
            if num_attr is not None:
                col_num = int(num_attr)

        if col_num == 2:
            print(f"PASS: Component 1 — 2-column layout detected (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected 2-column layout, found cols num={col_num}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Attendee names — ~24pt bold centered runs (0.30 points)
    # Golden has 120 runs with bold=True, size=304800 (24pt), paragraph centered.
    # Allow some tolerance on size: 22pt-26pt (279400 to 330200 EMU).
    try:
        name_runs = []
        for para in doc.paragraphs:
            is_centered = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            for run in para.runs:
                if (run.font.bold is True
                        and run.font.size is not None
                        and 279400 <= run.font.size <= 330200
                        and is_centered
                        and run.text.strip()):
                    name_runs.append(run.text.strip())

        name_count = len(name_runs)
        if name_count >= 115:
            # Allow small tolerance (115-120 acceptable for full credit)
            print(f"PASS: Component 2 — {name_count} attendee names in ~24pt bold centered (0.30 pts)")
            total_score += 0.30
        elif name_count >= 60:
            partial = 0.15
            print(f"PARTIAL: Component 2 — {name_count}/120 names found ({partial} pts)")
            total_score += partial
        elif name_count >= 10:
            partial = 0.05
            print(f"PARTIAL: Component 2 — {name_count}/120 names found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {name_count} attendee names in ~24pt bold centered found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Organization names — ~16pt centered runs (0.20 points)
    # Golden has 120 runs with size=203200 (16pt), centered, not bold.
    # Allow tolerance: 14pt-18pt (177800 to 228600 EMU).
    try:
        org_runs = []
        for para in doc.paragraphs:
            is_centered = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            for run in para.runs:
                if (run.font.size is not None
                        and 177800 <= run.font.size <= 228600
                        and is_centered
                        and run.text.strip()
                        and run.font.bold is not True):
                    org_runs.append(run.text.strip())

        org_count = len(org_runs)
        if org_count >= 115:
            print(f"PASS: Component 3 — {org_count} organization names in ~16pt centered (0.20 pts)")
            total_score += 0.20
        elif org_count >= 60:
            partial = 0.10
            print(f"PARTIAL: Component 3 — {org_count}/120 orgs found ({partial} pts)")
            total_score += partial
        elif org_count >= 10:
            partial = 0.03
            print(f"PARTIAL: Component 3 — {org_count}/120 orgs found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {org_count} organization names in ~16pt centered found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Page/column breaks for badge separation (0.15 points)
    # Golden has 59 page breaks + 60 column breaks. Need substantial breaks for 120 badges.
    try:
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        ns = {'w': ns_w}
        page_breaks = 0
        column_breaks = 0
        for para in doc.paragraphs:
            for run in para.runs:
                for br in run.element.findall('.//w:br', ns):
                    btype = br.attrib.get(f'{{{ns_w}}}type', '')
                    if btype == 'page':
                        page_breaks += 1
                    elif btype == 'column':
                        column_breaks += 1

        total_breaks = page_breaks + column_breaks
        if total_breaks >= 100:
            print(f"PASS: Component 4 — {page_breaks} page breaks + {column_breaks} column breaks (0.15 pts)")
            total_score += 0.15
        elif total_breaks >= 50:
            partial = 0.08
            print(f"PARTIAL: Component 4 — {total_breaks} total breaks ({partial} pts)")
            total_score += partial
        elif total_breaks >= 10:
            partial = 0.03
            print(f"PARTIAL: Component 4 — {total_breaks} total breaks ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {total_breaks} total breaks (page+column), expected ~119")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Badge structure — name+org pairs follow each other (0.15 points)
    # Each badge should have a name paragraph immediately followed by an org paragraph.
    # Check that there are ~120 such consecutive pairs.
    try:
        pair_count = 0
        paragraphs = doc.paragraphs
        for i in range(len(paragraphs) - 1):
            p_name = paragraphs[i]
            p_org = paragraphs[i + 1]

            # Name paragraph: centered, has a bold ~24pt run
            name_ok = False
            if p_name.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                for r in p_name.runs:
                    if (r.font.bold is True
                            and r.font.size is not None
                            and 279400 <= r.font.size <= 330200
                            and r.text.strip()):
                        name_ok = True
                        break

            # Org paragraph: centered, has a ~16pt non-bold run
            org_ok = False
            if name_ok and p_org.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                for r in p_org.runs:
                    if (r.font.size is not None
                            and 177800 <= r.font.size <= 228600
                            and r.text.strip()
                            and r.font.bold is not True):
                        org_ok = True
                        break

            if name_ok and org_ok:
                pair_count += 1

        if pair_count >= 115:
            print(f"PASS: Component 5 — {pair_count} name+org pairs in correct structure (0.15 pts)")
            total_score += 0.15
        elif pair_count >= 60:
            partial = 0.08
            print(f"PARTIAL: Component 5 — {pair_count}/120 pairs ({partial} pts)")
            total_score += partial
        elif pair_count >= 10:
            partial = 0.03
            print(f"PARTIAL: Component 5 — {pair_count}/120 pairs ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {pair_count} name+org pairs found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
