"""
Initial Setup: Magazine article about Istanbul's Ottoman-era gardens (no drop caps)
Task ID: writer_para_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_para_051'
OUTPUT = f'{WORKDIR}/magazine_article.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1 - Article title
    heading1 = doc.add_heading('The Forgotten Gardens of Istanbul', level=1)

    # Paragraph 2: Heading 2 - Part One
    heading2 = doc.add_heading('Part One: Discovery', level=2)

    # Paragraph 3: First body paragraph of Part One (NO drop cap in initial)
    para3 = doc.add_paragraph(
        'Behind the bustling streets of Sultanahmet, hidden from the tourist crowds '
        'and guidebook itineraries, lies a network of Ottoman-era gardens that have '
        'survived centuries of urban transformation.'
    )

    # Paragraph 4: Second body paragraph of Part One
    para4 = doc.add_paragraph(
        'These gardens, known locally as the \u201cgreen rooms,\u201d were originally designed '
        'as private retreats for the city elite. Today, fewer than a dozen remain intact.'
    )

    # Paragraph 5: Heading 2 - Part Two
    heading5 = doc.add_heading('Part Two: Restoration', level=2)

    # Paragraph 6: First body paragraph of Part Two (NO drop cap in initial)
    para6 = doc.add_paragraph(
        'In 2018, a group of landscape architects and historians launched an ambitious '
        'project to restore five of these gardens to their original splendor.'
    )

    # Paragraph 7: Second body paragraph of Part Two
    para7 = doc.add_paragraph(
        'The restoration work uncovered an irrigation system dating back to the 16th century, '
        'remarkably similar to Persian garden engineering techniques.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
