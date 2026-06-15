"""
Reward Script: Set header row background to dark green and text to white
Task ID: writer_tm_013
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): All 4 header cells have dark green background (#006400)
  Component 2 (0.4): All header cell runs with text have white font color (#FFFFFF)
  Gate: Data rows must remain unchanged (no green bg applied to data rows)
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_013'


def color_distance_rgb(c1_tuple, c2_tuple):
    """Euclidean distance between two RGB tuples."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1_tuple, c2_tuple)))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least one table with at least 2 rows
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) < 2:
        print("FAIL: Table has fewer than 2 rows")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: Header row background color is dark green #006400 (0.6 points) ---
    try:
        header_row = table.rows[0]
        green_bg_count = 0
        total_header_cells = len(header_row.cells)

        for ci, cell in enumerate(header_row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill is not None:
                        fill_upper = fill.upper().lstrip('#')
                        # Check if it's close to 006400 (dark green)
                        try:
                            r = int(fill_upper[0:2], 16)
                            g = int(fill_upper[2:4], 16)
                            b = int(fill_upper[4:6], 16)
                            dist = color_distance_rgb((r, g, b), (0, 100, 0))
                            if dist < 30:  # tolerance for close shades
                                green_bg_count += 1
                                print(f"  PASS: Cell [0,{ci}] bg=#{fill_upper} (dist={dist:.1f})")
                            else:
                                print(f"  FAIL: Cell [0,{ci}] bg=#{fill_upper} not close to #006400 (dist={dist:.1f})")
                        except (ValueError, IndexError):
                            print(f"  FAIL: Cell [0,{ci}] invalid fill color: {fill}")
                    else:
                        print(f"  FAIL: Cell [0,{ci}] no fill attribute in shading")
                else:
                    print(f"  FAIL: Cell [0,{ci}] no shading element")
            else:
                print(f"  FAIL: Cell [0,{ci}] no tcPr element")

        if green_bg_count == total_header_cells and total_header_cells > 0:
            print(f"PASS: Component 1 - All {green_bg_count}/{total_header_cells} header cells have dark green background (0.6 pts)")
            total_score += 0.6
        elif green_bg_count > 0:
            partial = 0.6 * (green_bg_count / total_header_cells)
            print(f"PARTIAL: Component 1 - {green_bg_count}/{total_header_cells} header cells have dark green background ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - No header cells have dark green background")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # --- Component 2: Header row text is white #FFFFFF (0.4 points) ---
    try:
        header_row = table.rows[0]
        white_text_cells = 0
        total_header_cells = len(header_row.cells)

        for ci, cell in enumerate(header_row.cells):
            cell_has_white = True
            runs_with_text = 0
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():  # only check runs that have text
                        runs_with_text += 1
                        rgb = run.font.color.rgb
                        if rgb is not None:
                            r, g, b = rgb[0], rgb[1], rgb[2]
                            dist = color_distance_rgb((r, g, b), (255, 255, 255))
                            if dist > 30:
                                cell_has_white = False
                                print(f"  FAIL: Cell [0,{ci}] run text='{run.text}' color={rgb} not white (dist={dist:.1f})")
                        else:
                            # None means inherited - likely black (default), not white
                            cell_has_white = False
                            print(f"  FAIL: Cell [0,{ci}] run text='{run.text}' color=None (inherited, likely not white)")

            if runs_with_text > 0 and cell_has_white:
                white_text_cells += 1
                print(f"  PASS: Cell [0,{ci}] all text runs are white")

        if white_text_cells == total_header_cells and total_header_cells > 0:
            print(f"PASS: Component 2 - All {white_text_cells}/{total_header_cells} header cells have white text (0.4 pts)")
            total_score += 0.4
        elif white_text_cells > 0:
            partial = 0.4 * (white_text_cells / total_header_cells)
            print(f"PARTIAL: Component 2 - {white_text_cells}/{total_header_cells} header cells have white text ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - No header cells have white text")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # --- Gate: Data rows must remain unchanged (no green bg applied to data rows) ---
    # This is a precondition gate, not a scoring component. If data rows were
    # incorrectly modified, we penalize by zeroing out the score.
    try:
        data_rows_ok = True
        num_data_rows = len(table.rows) - 1

        for ri in range(1, len(table.rows)):
            row = table.rows[ri]
            for ci, cell in enumerate(row.cells):
                # Check no green background was applied to data rows
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill is not None:
                            fill_upper = fill.upper().lstrip('#')
                            try:
                                r = int(fill_upper[0:2], 16)
                                g = int(fill_upper[2:4], 16)
                                b = int(fill_upper[4:6], 16)
                                dist = color_distance_rgb((r, g, b), (0, 100, 0))
                                if dist < 30:
                                    data_rows_ok = False
                                    print(f"  FAIL: Data cell [{ri},{ci}] has green background (should be unchanged)")
                            except (ValueError, IndexError):
                                pass

                # Check no white text on runs with actual text content in data rows
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip() and run.font.color.rgb is not None:
                            r2, g2, b2 = run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2]
                            dist = color_distance_rgb((r2, g2, b2), (255, 255, 255))
                            if dist < 30:
                                data_rows_ok = False
                                print(f"  FAIL: Data cell [{ri},{ci}] has white text (should be unchanged)")

        if data_rows_ok:
            print(f"GATE PASS: All {num_data_rows} data rows remain unchanged")
        else:
            print(f"GATE FAIL: Some data rows were modified - zeroing score")
            total_score = 0.0
    except Exception as e:
        print(f"ERROR: Data rows gate check - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
