"""
Initial Setup: Set header row background to dark green and text to white in pricing table
Task ID: writer_tm_013
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_013'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Pricing Sheet', level=1)

    # Add a brief intro paragraph
    doc.add_paragraph(
        'Below is the current pricing for our Q2 2025 product catalog. '
        'Please review and confirm the totals before submitting to procurement.'
    )

    # Create the 4x6 pricing table (1 header + 5 data rows)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # Row 1: Headers (plain black text, white/default background)
    headers = ['Item', 'Unit Price', 'Quantity', 'Total']
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black text

    # Data rows with realistic content
    data = [
        ['Wireless Keyboard', '$45.99', '12', '$551.88'],
        ['USB-C Hub Adapter', '$32.50', '8', '$260.00'],
        ['Ergonomic Mouse Pad', '$18.75', '25', '$468.75'],
        ['Monitor Stand Riser', '$67.00', '6', '$402.00'],
        ['Cable Management Kit', '$14.25', '15', '$213.75'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black text

    # Add a closing paragraph
    doc.add_paragraph('')
    doc.add_paragraph(
        'Note: All prices are in USD and subject to bulk discount negotiations.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
