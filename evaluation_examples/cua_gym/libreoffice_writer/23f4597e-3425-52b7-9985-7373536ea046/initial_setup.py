"""
Initial Setup: Insert footnotes in Introduction paragraph
Task ID: writer_bs_017
Domain: libreoffice_writer

Creates a research paper document with an Introduction paragraph containing
'machine learning', 'neural networks', and 'deep learning'. No footnotes.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_017'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('Advances in Computational Intelligence: A Survey', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Authors ---
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_a = authors.add_run('Dr. Elena Vasquez, Prof. Rajesh Patel, Dr. Mei-Ling Zhou')
    run_a.font.size = Pt(11)
    run_a.italic = True

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_af = affil.add_run('Department of Computer Science, Stanford University')
    run_af.font.size = Pt(10)
    run_af.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        'This survey examines the rapid evolution of computational intelligence techniques '
        'over the past two decades. We review major algorithmic advances, benchmark performance '
        'improvements, and emerging application domains. Our analysis spans supervised and '
        'unsupervised methods, optimization frameworks, and hybrid architectures that integrate '
        'symbolic reasoning with data-driven approaches. We identify key open problems and '
        'propose directions for future research in scalable, interpretable, and robust systems.'
    )
    p_abs = doc.add_paragraph(abstract_text)
    p_abs.paragraph_format.space_after = Pt(12)

    # --- 1. Introduction (KEY paragraph with target phrases) ---
    doc.add_heading('1. Introduction', level=1)
    intro_text = (
        'The field of artificial intelligence has undergone a remarkable transformation '
        'in recent years. At its core, machine learning provides algorithms that improve '
        'automatically through experience, enabling systems to discover patterns in vast '
        'datasets without explicit programming. Building on these foundations, neural networks '
        'offer a biologically inspired computational paradigm that has proven especially '
        'effective for tasks involving perception and pattern recognition. More recently, '
        'deep learning has emerged as a powerful extension of these ideas, leveraging '
        'multi-layered architectures to learn hierarchical representations of data. Together, '
        'these approaches have driven breakthroughs in computer vision, natural language '
        'processing, robotics, and scientific discovery.'
    )
    p_intro = doc.add_paragraph(intro_text)
    p_intro.paragraph_format.space_after = Pt(6)

    # --- 2. Literature Review ---
    doc.add_heading('2. Literature Review', level=1)
    lit_text = (
        'Early work in computational intelligence drew heavily from statistical learning theory '
        'and Bayesian inference. Vapnik and Chervonenkis established foundational bounds on '
        'generalization error that remain central to the field. The resurgence of connectionist '
        'models in the 1980s, driven by the backpropagation algorithm, opened new avenues for '
        'practical applications. By the mid-2000s, kernel methods and ensemble techniques such '
        'as random forests and gradient boosting had become standard tools in the practitioner\'s '
        'toolkit. The ImageNet competition of 2012 marked a watershed moment, demonstrating '
        'that convolutional architectures could dramatically outperform traditional feature '
        'engineering approaches on large-scale visual recognition tasks.'
    )
    doc.add_paragraph(lit_text)

    # --- 3. Methodology ---
    doc.add_heading('3. Methodology', level=1)
    meth_text = (
        'Our survey methodology follows a systematic review protocol. We queried IEEE Xplore, '
        'ACM Digital Library, and arXiv for papers published between 2015 and 2025 using '
        'predefined search terms. After de-duplication, we screened 4,327 candidate papers by '
        'title and abstract, yielding 892 papers for full-text review. We categorized these by '
        'application domain, algorithmic family, and reported performance metrics. Inter-rater '
        'reliability was measured using Cohen\'s kappa, achieving a coefficient of 0.87 across '
        'all classification dimensions.'
    )
    doc.add_paragraph(meth_text)

    # --- 4. Results ---
    doc.add_heading('4. Results', level=1)
    results_text = (
        'Our analysis reveals several clear trends. First, transformer-based architectures '
        'have increasingly dominated performance benchmarks across modalities, from text to '
        'images to audio. Second, self-supervised pre-training followed by task-specific '
        'fine-tuning has become the prevailing paradigm, reducing the need for large labeled '
        'datasets. Third, there is growing interest in efficiency-oriented research, including '
        'model compression, knowledge distillation, and hardware-aware neural architecture '
        'search. Finally, interpretability and fairness have received increased attention, '
        'with new metrics and frameworks proposed for evaluating model transparency.'
    )
    doc.add_paragraph(results_text)

    # --- Add a simple table for survey statistics ---
    doc.add_heading('4.1 Survey Statistics', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Category', 'Papers Reviewed', 'Key Trend']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data = [
        ['Computer Vision', '247', 'Vision Transformers surpass CNNs'],
        ['NLP', '312', 'Large language models dominate'],
        ['Reinforcement Learning', '189', 'Sim-to-real transfer improves'],
        ['Generative Models', '144', 'Diffusion models gain traction'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- 5. Conclusion ---
    doc.add_heading('5. Conclusion', level=1)
    conc_text = (
        'Computational intelligence continues to advance at an accelerating pace. The '
        'convergence of increased computational resources, large-scale datasets, and algorithmic '
        'innovations has enabled applications that were considered intractable only a decade ago. '
        'Moving forward, we anticipate that research will increasingly focus on building systems '
        'that are not only accurate but also efficient, interpretable, and aligned with human '
        'values. Interdisciplinary collaboration between computer scientists, domain experts, '
        'and ethicists will be essential to realizing the full potential of these technologies.'
    )
    doc.add_paragraph(conc_text)

    # --- References ---
    doc.add_heading('References', level=1)
    refs = [
        'Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.',
        'LeCun, Y., Bottou, L., Bengio, Y., & Haffner, P. (1998). Gradient-based learning applied to document recognition. Proceedings of the IEEE, 86(11), 2278-2324.',
        'Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention is all you need. NeurIPS 2017.',
        'He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. CVPR 2016.',
        'Devlin, J., Chang, M., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers. NAACL 2019.',
    ]
    for ref in refs:
        p_ref = doc.add_paragraph(ref, style='List Number')
        p_ref.paragraph_format.space_after = Pt(2)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
