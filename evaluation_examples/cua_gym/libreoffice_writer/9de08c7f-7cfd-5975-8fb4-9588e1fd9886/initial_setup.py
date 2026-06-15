"""
Initial Setup: Create a Writer document with a References section (no special formatting)
Task ID: writer_fs_080
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_080'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Impact of Urbanization on Regional Biodiversity', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Authors ---
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Dr. Elena Vasquez, Prof. James Whitfield, Dr. Anika Patel')
    run.font.size = Pt(11)
    run.font.italic = True

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil.add_run('Department of Environmental Sciences, Lakewood University')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This paper examines the relationship between urban expansion and biodiversity '
        'loss across 15 metropolitan regions in North America. Using satellite imagery '
        'from 2005 to 2024, we quantify habitat fragmentation and correlate it with '
        'species richness data from field surveys. Our findings indicate a 23% decline '
        'in native plant diversity and a 31% reduction in pollinator populations in '
        'areas experiencing rapid suburban development. We propose a set of mitigation '
        'strategies including wildlife corridors, green infrastructure mandates, and '
        'urban rewilding programs that could reverse current trends.'
    )

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Urbanization is among the most significant drivers of global biodiversity loss. '
        'As cities expand, natural habitats are converted into residential, commercial, and '
        'industrial zones, leading to habitat fragmentation, altered hydrology, and increased '
        'pollution. The United Nations projects that 68% of the world\'s population will reside '
        'in urban areas by 2050, intensifying pressure on peri-urban ecosystems.'
    )
    doc.add_paragraph(
        'Previous studies have largely focused on single-city case analyses, making it '
        'difficult to generalize findings across different biogeographic regions. This study '
        'addresses that gap by employing a multi-city, multi-decade approach to quantify '
        'how different rates and patterns of urbanization affect species richness, community '
        'composition, and ecosystem services.'
    )

    # --- Methodology ---
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        'We selected 15 metropolitan regions representing diverse climatic zones: '
        'Seattle, Denver, Minneapolis, Atlanta, Houston, Phoenix, Chicago, Toronto, '
        'Portland, San Diego, Nashville, Boston, Miami, Dallas, and Charlotte. For each '
        'region, we obtained Landsat and Sentinel-2 satellite imagery at five-year intervals '
        'from 2005 to 2024.'
    )
    doc.add_paragraph(
        'Land cover classification was performed using a random forest algorithm with '
        'an overall accuracy of 92.4%. Field surveys were conducted in collaboration with '
        'local naturalist networks, recording presence/absence data for vascular plants, '
        'Lepidoptera, Hymenoptera, and avian species within standardized 1-hectare plots.'
    )

    # --- Results ---
    doc.add_heading('3. Results', level=1)
    doc.add_paragraph(
        'Across all 15 regions, impervious surface area increased by an average of 18.7% '
        'between 2005 and 2024. Native plant species richness declined by 23.1% in zones '
        'where impervious cover exceeded 60%. Pollinator abundance showed a steeper decline '
        'of 31.4%, with solitary bee populations most severely affected.'
    )
    doc.add_paragraph(
        'Cities with established green infrastructure programs (Portland, Seattle, Toronto) '
        'exhibited significantly lower biodiversity loss (p < 0.01) compared to cities without '
        'such programs. Wildlife corridors connecting fragmented patches were associated with '
        'a 14% higher species retention rate.'
    )

    # --- Discussion ---
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph(
        'Our results confirm that urbanization is a primary driver of local biodiversity '
        'decline, but the magnitude of loss is strongly modulated by planning interventions. '
        'Cities that incorporated ecological connectivity into their master plans maintained '
        'higher ecosystem resilience. The effectiveness of green corridors was particularly '
        'notable for mobile taxa such as birds and butterflies.'
    )
    doc.add_paragraph(
        'We recommend that municipalities adopt three complementary strategies: (1) mandating '
        'minimum green space ratios in new developments, (2) retrofitting existing corridors '
        'with native plant communities, and (3) implementing pollinator-friendly landscaping '
        'ordinances. These measures have demonstrated measurable benefits within 3-5 years '
        'of implementation in our study regions.'
    )

    # --- References (NO special formatting - plain section) ---
    doc.add_heading('References', level=1)
    references = [
        'Anderson, K. & Park, S. (2018). Urban heat islands and their effect on '
        'arthropod communities. Journal of Urban Ecology, 14(2), 112-128.',

        'Beninde, J., Veith, M., & Hochkirch, A. (2015). Biodiversity in cities needs '
        'space: a meta-analysis of factors determining intra-urban biodiversity variation. '
        'Ecology Letters, 18(6), 581-592.',

        'Concepcion, E.D., Moretti, M., Altermatt, F., Nobis, M.P., & Obrist, M.K. '
        '(2015). Impacts of urbanisation on biodiversity: the role of species mobility, '
        'degree of specialisation and spatial scale. Oikos, 124(12), 1571-1582.',

        'Goddard, M.A., Dougill, A.J., & Benton, T.G. (2010). Scaling up from gardens: '
        'biodiversity conservation in urban environments. Trends in Ecology & Evolution, '
        '25(2), 90-98.',

        'Hall, D.M., Camilo, G.R., Tonietto, R.K., Ollerton, J., Ahrne, K., Arduser, M., '
        'Ascher, J.S., Baldock, K.C., & Fowler, R. (2017). The city as a refuge for '
        'insect pollinators. Conservation Biology, 31(1), 24-29.',

        'Ives, C.D., Lentini, P.E., Threlfall, C.G., Ikin, K., Shanahan, D.F., Garrard, G.E., '
        '& Bekessy, S.A. (2016). Cities are hotspots for threatened species. Global Ecology '
        'and Biogeography, 25(1), 117-126.',

        'McKinney, M.L. (2008). Effects of urbanization on species richness: a review of '
        'plants and animals. Urban Ecosystems, 11(2), 161-176.',

        'Newbold, T., Hudson, L.N., Hill, S.L., Contu, S., Lysenko, I., Senior, R.A., '
        '& Purvis, A. (2015). Global effects of land use on local terrestrial biodiversity. '
        'Nature, 520(7545), 45-50.',

        'Seto, K.C., Guneralp, B., & Hutyra, L.R. (2012). Global forecasts of urban expansion '
        'to 2030 and direct impacts on biodiversity and carbon pools. Proceedings of the '
        'National Academy of Sciences, 109(40), 16083-16088.',

        'Threlfall, C.G., Walker, K., Williams, N.S., Hahs, A.K., Mata, L., Stork, N., & '
        'Livesley, S.J. (2015). The conservation value of urban green space habitats for '
        'Australian native bee communities. Biological Conservation, 187, 240-248.',
    ]

    for ref in references:
        doc.add_paragraph(ref)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
