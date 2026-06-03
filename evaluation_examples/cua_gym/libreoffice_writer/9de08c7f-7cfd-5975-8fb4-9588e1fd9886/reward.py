"""
Reward Script: Set References section background color and indentation
Task ID: writer_fs_080
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Background color #FFFACD on all References paragraphs
  Component 2 (0.4): Left and right indent ~1 cm on all References paragraphs
  Component 3 (0.2): Non-References paragraphs remain unaffected (no shading/indent)
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_080'

# 1 cm in EMU = 360000; allow +/- 10% tolerance
ONE_CM_EMU = 360000
INDENT_TOLERANCE = 36000  # 10% of 1 cm


def find_references_range(doc):
    """
    Find the paragraph index of 'References' heading and all subsequent
    paragraphs (the References section extends to end of document).
    Returns (start_idx, end_idx) inclusive, or (None, None) if not found.
    """
    start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'References' and 'Heading' in para.style.name:
            start = i
            break
    if start is None:
        return None, None
    return start, len(doc.paragraphs) - 1


def get_shading_fill(para):
    """Extract the w:fill attribute from paragraph shading, or None."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    pPr = para._element.find(f'{{{ns}}}pPr')
    if pPr is None:
        return None
    shd = pPr.find(f'{{{ns}}}shd')
    if shd is None:
        return None
    return shd.get(qn('w:fill'))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: find the References section
    ref_start, ref_end = find_references_range(doc)
    if ref_start is None:
        print("CRITICAL: No 'References' heading found in document")
        print("REWARD: 0.0")
        return 0.0

    ref_paras = list(range(ref_start, ref_end + 1))
    non_ref_paras = list(range(0, ref_start))
    num_ref = len(ref_paras)

    print(f"INFO: References section spans paragraphs {ref_start}-{ref_end} ({num_ref} paragraphs)")

    # Component 1: Background color #FFFACD on all References paragraphs (0.4 points)
    try:
        shading_pass = 0
        for idx in ref_paras:
            para = doc.paragraphs[idx]
            fill = get_shading_fill(para)
            if fill is not None and fill.upper() == 'FFFACD':
                shading_pass += 1
            else:
                print(f"FAIL: Para {idx} shading fill={fill}, expected FFFACD")

        if shading_pass == num_ref:
            print(f"PASS: Component 1 -- All {num_ref} References paragraphs have #FFFACD background (0.4 pts)")
            total_score += 0.4
        elif shading_pass > 0:
            partial = 0.4 * (shading_pass / num_ref)
            print(f"PARTIAL: Component 1 -- {shading_pass}/{num_ref} paragraphs have correct shading ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- No References paragraphs have #FFFACD background")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Left and right indent ~1 cm on all References paragraphs (0.4 points)
    try:
        indent_pass = 0
        for idx in ref_paras:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            left = pf.left_indent
            right = pf.right_indent

            if left is None or right is None:
                print(f"FAIL: Para {idx} indent left={left}, right={right} (expected ~{ONE_CM_EMU})")
                continue

            left_val = int(left)
            right_val = int(right)
            left_ok = abs(left_val - ONE_CM_EMU) <= INDENT_TOLERANCE
            right_ok = abs(right_val - ONE_CM_EMU) <= INDENT_TOLERANCE

            if left_ok and right_ok:
                indent_pass += 1
            else:
                print(f"FAIL: Para {idx} indent left={left_val}, right={right_val} (expected ~{ONE_CM_EMU})")

        if indent_pass == num_ref:
            print(f"PASS: Component 2 -- All {num_ref} References paragraphs have ~1cm left+right indent (0.4 pts)")
            total_score += 0.4
        elif indent_pass > 0:
            partial = 0.4 * (indent_pass / num_ref)
            print(f"PARTIAL: Component 2 -- {indent_pass}/{num_ref} paragraphs have correct indent ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- No References paragraphs have ~1cm indent")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Formatting is section-specific (0.2 points)
    # References paragraphs HAVE the formatting AND non-References paragraphs do NOT.
    # This is a compound check anchored to the task change: only passes when
    # References formatting exists AND is correctly scoped.
    try:
        # First gate: count References paragraphs that have the formatting
        ref_formatted_count = sum(
            1 for idx in ref_paras
            if get_shading_fill(doc.paragraphs[idx]) is not None
            and get_shading_fill(doc.paragraphs[idx]).upper() == 'FFFACD'
        )

        if ref_formatted_count == 0:
            print("FAIL: Component 3 -- No References paragraphs have formatting, so scope check is moot")
        else:
            contaminated = 0
            for idx in non_ref_paras:
                para = doc.paragraphs[idx]
                fill = get_shading_fill(para)
                pf = para.paragraph_format

                has_ref_shading = fill is not None and fill.upper() == 'FFFACD'
                has_ref_indent = (
                    pf.left_indent is not None
                    and pf.right_indent is not None
                    and abs(int(pf.left_indent) - ONE_CM_EMU) <= INDENT_TOLERANCE
                    and abs(int(pf.right_indent) - ONE_CM_EMU) <= INDENT_TOLERANCE
                )

                if has_ref_shading or has_ref_indent:
                    contaminated += 1
                    print(f"FAIL: Non-ref para {idx} has References formatting")

            if contaminated == 0:
                print(f"PASS: Component 3 -- Formatting correctly scoped to References section only (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 -- {contaminated} non-References paragraphs have References formatting")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
