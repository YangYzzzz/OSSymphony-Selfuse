"""
Initial Setup: Enable widow/orphan control and 'Do not split paragraph' for Quotations style
Task ID: writer_fs_048
Domain: libreoffice_writer

Creates a 30-page Writer document with realistic content. Several paragraphs use
the 'Quotations' style (block quotes). Default Paragraph Style has NO widow/orphan
control. Quotations style does NOT have 'keep together' set.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# --- Realistic long-form content ---

INTRO_TEXT = (
    "The rapid advancement of artificial intelligence technologies has fundamentally "
    "transformed how organizations approach data analysis, customer engagement, and "
    "operational efficiency. Companies across diverse sectors are discovering that "
    "machine learning algorithms can process vast amounts of information far more "
    "quickly than traditional methods, enabling decision-makers to act on insights "
    "that would have been impossible to extract just a decade ago."
)

BODY_PARAGRAPHS = [
    (
        "In the healthcare sector, predictive analytics models are being deployed to "
        "identify patients at high risk of developing chronic conditions. Researchers at "
        "the Stanford Medical Center reported a 34% improvement in early detection rates "
        "for cardiovascular disease when AI-assisted screening protocols were introduced "
        "in their outpatient clinics during the 2024 fiscal year. These results have "
        "prompted similar institutions to invest heavily in computational infrastructure."
    ),
    (
        "Meanwhile, the financial services industry continues to leverage natural "
        "language processing for regulatory compliance. Banks such as Deutsche Bank and "
        "JPMorgan Chase have automated the review of millions of internal communications "
        "to flag potential compliance violations. According to a 2025 report by McKinsey "
        "& Company, these automation efforts have reduced manual review costs by an "
        "estimated $2.8 billion annually across the top 20 global financial institutions."
    ),
    (
        "The manufacturing sector has embraced computer vision systems for quality "
        "control on production lines. Toyota's Nagoya plant implemented an AI-based "
        "visual inspection system in March 2024 that examines each component at a rate "
        "of 1,200 units per minute, identifying defects with 99.7% accuracy. The system "
        "replaced a team of 45 human inspectors and reduced defect-related recalls by "
        "62% in the first six months of operation."
    ),
    (
        "Education technology platforms have adopted adaptive learning algorithms that "
        "personalize curriculum delivery based on individual student performance metrics. "
        "Coursera's adaptive engine, launched in September 2024, adjusts the difficulty "
        "and pacing of course material in real time, resulting in a 28% increase in "
        "course completion rates and a 15-point improvement in average assessment scores "
        "among students enrolled in data science certificate programs."
    ),
    (
        "The agricultural industry has seen significant gains from precision farming "
        "techniques powered by satellite imagery and IoT sensor networks. In Brazil's "
        "Mato Grosso state, soybean farmers using AI-optimized irrigation schedules "
        "reported yield increases of 18% while reducing water consumption by 23% during "
        "the 2024-2025 growing season. These results have attracted $450 million in "
        "venture capital funding for agricultural AI startups in Latin America alone."
    ),
    (
        "Retail organizations have deployed recommendation engines that analyze purchase "
        "history, browsing behavior, and demographic data to generate personalized product "
        "suggestions. Amazon's latest recommendation model, updated in January 2025, "
        "accounts for 38% of all product purchases on the platform, up from 35% the "
        "previous year. Walmart and Target have similarly invested in machine learning "
        "infrastructure to compete in the personalization space."
    ),
    (
        "The energy sector has applied reinforcement learning algorithms to optimize "
        "power grid management. The European Network of Transmission System Operators "
        "reported that AI-driven load balancing reduced energy waste by 12% across the "
        "continental grid in 2024. Wind farm operators in Denmark have used predictive "
        "models to anticipate turbine maintenance needs, reducing unplanned downtime by "
        "41% and saving approximately EUR 180 million in repair costs."
    ),
    (
        "In the legal profession, large language models have been integrated into "
        "document review workflows for litigation and due diligence. Firms such as "
        "Baker McKenzie and Clifford Chance now use AI assistants to summarize case law, "
        "draft preliminary contract clauses, and identify relevant precedents. A survey "
        "by the American Bar Association found that 67% of large law firms had adopted "
        "some form of AI-assisted research tool by the end of 2024."
    ),
    (
        "Transportation and logistics companies have implemented route optimization "
        "algorithms that account for real-time traffic conditions, weather forecasts, "
        "and delivery time windows. UPS's ORION system, enhanced with deep learning "
        "capabilities in 2024, now saves the company approximately 100 million miles "
        "of driving per year, translating to fuel savings of over $400 million and a "
        "reduction of 100,000 metric tons of carbon emissions annually."
    ),
    (
        "The cybersecurity landscape has been reshaped by anomaly detection systems "
        "that identify potential threats in network traffic patterns. CrowdStrike's "
        "Falcon platform processes over 2 trillion security events per week, using "
        "graph neural networks to correlate seemingly unrelated indicators of compromise. "
        "In 2024, the platform detected and prevented an estimated 4.2 million attempted "
        "breaches across its client base of 23,000 organizations worldwide."
    ),
]

QUOTE_BLOCKS = [
    (
        "As noted in the 2025 Global AI Index published by Stanford University's "
        "Institute for Human-Centered Artificial Intelligence, 'The convergence of "
        "large language models, multimodal architectures, and retrieval-augmented "
        "generation has created a new paradigm for enterprise knowledge management. "
        "Organizations that fail to adopt these technologies risk falling behind "
        "competitors who can extract actionable intelligence from their data assets "
        "at unprecedented speed and scale. The economic implications are staggering: "
        "our models project that AI-driven productivity gains will contribute an "
        "additional $15.7 trillion to global GDP by 2030, with the largest gains "
        "concentrated in North America and East Asia.'"
    ),
    (
        "The World Economic Forum's Future of Jobs Report 2025 states: 'While "
        "automation will displace an estimated 85 million jobs globally by 2027, "
        "it will simultaneously create 97 million new roles that are better adapted "
        "to the new division of labor between humans, machines, and algorithms. "
        "The net positive effect of 12 million jobs, however, masks significant "
        "regional and sectoral disparities that policymakers must address through "
        "targeted reskilling programs and social safety nets. Countries that invest "
        "in AI literacy at the primary education level will see the most balanced "
        "labor market transitions over the coming decade.'"
    ),
    (
        "According to Dr. Elena Rodriguez, Chief Technology Officer at Siemens "
        "Digital Industries: 'The integration of digital twin technology with "
        "generative AI has opened entirely new possibilities for industrial design "
        "and process optimization. Our engineers can now simulate thousands of "
        "manufacturing configurations in hours rather than weeks, identifying "
        "optimal parameters that reduce material waste by up to 30% while "
        "maintaining or improving product quality. This represents a fundamental "
        "shift in how we approach engineering challenges, moving from intuition-based "
        "design to data-driven exploration of the solution space.'"
    ),
    (
        "Professor James Chen of MIT's Computer Science and Artificial Intelligence "
        "Laboratory observed in his keynote address at NeurIPS 2024: 'The emergence "
        "of reasoning-capable language models has blurred the traditional boundary "
        "between narrow and general artificial intelligence. While we remain far "
        "from true AGI, the ability of current systems to decompose complex problems, "
        "apply multi-step logical reasoning, and generate novel hypotheses suggests "
        "that the next five years will bring capabilities that most researchers "
        "considered decades away just three years ago.'"
    ),
    (
        "The International Monetary Fund's April 2025 World Economic Outlook "
        "highlighted: 'Central banks face a new challenge as AI-driven productivity "
        "gains create disinflationary pressures in advanced economies while "
        "simultaneously increasing demand for energy and semiconductor inputs. "
        "The resulting supply-demand dynamics in commodity markets add a layer "
        "of complexity to monetary policy decisions that traditional macroeconomic "
        "models are ill-equipped to capture. We recommend that central banks invest "
        "in AI-augmented forecasting tools to better anticipate these novel "
        "transmission channels.'"
    ),
]

CLOSING_PARAGRAPHS = [
    (
        "Looking ahead, the trajectory of artificial intelligence development "
        "suggests that the most transformative applications are yet to come. "
        "Multimodal models capable of processing text, images, audio, and video "
        "simultaneously are opening new frontiers in content creation, scientific "
        "research, and human-computer interaction. The convergence of AI with "
        "quantum computing, biotechnology, and materials science promises to "
        "accelerate innovation across all domains of human endeavor."
    ),
    (
        "However, these advances also raise important ethical considerations "
        "regarding privacy, algorithmic bias, and the concentration of technological "
        "power. Regulatory frameworks such as the European Union's AI Act, enacted "
        "in 2024, represent the first comprehensive attempt to govern AI deployment "
        "at scale. Similar legislation is under consideration in the United States, "
        "China, and India, reflecting a global consensus that the benefits of AI "
        "must be balanced against potential risks to individual rights and social "
        "cohesion."
    ),
    (
        "The challenge for organizations is not merely to adopt AI technologies "
        "but to integrate them thoughtfully into existing workflows while maintaining "
        "human oversight and accountability. Companies that succeed in this endeavor "
        "will be those that invest not only in technical infrastructure but also in "
        "the training and empowerment of their workforce, ensuring that AI serves "
        "as a tool for human augmentation rather than replacement."
    ),
]


def create_initial():
    doc = Document()

    # --- Set default paragraph style: explicitly NO widow/orphan control ---
    default_style = doc.styles['Normal']
    pPr = default_style.element.get_or_add_pPr()
    # Ensure widow/orphan are explicitly 0 (disabled)
    widow_el = pPr.find(qn('w:widowControl'))
    if widow_el is not None:
        pPr.remove(widow_el)
    wc = parse_xml(f'<w:widowControl {nsdecls("w")} w:val="0"/>')
    pPr.append(wc)

    # --- Create a 'Quotations' style (block quote) without keep_together ---
    from docx.enum.style import WD_STYLE_TYPE
    quote_style = doc.styles.add_style('Quotations', WD_STYLE_TYPE.PARAGRAPH)
    quote_style.base_style = doc.styles['Normal']
    quote_fmt = quote_style.paragraph_format
    quote_fmt.left_indent = Inches(0.75)
    quote_fmt.right_indent = Inches(0.75)
    quote_fmt.space_before = Pt(12)
    quote_fmt.space_after = Pt(12)
    quote_fmt.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Explicitly do NOT set keep_together -- ensure it's off
    quote_pPr = quote_style.element.get_or_add_pPr()
    kt_el = quote_pPr.find(qn('w:keepLines'))
    if kt_el is not None:
        quote_pPr.remove(kt_el)
    # Set italic font for the quote style
    quote_font = quote_style.font
    quote_font.italic = True
    quote_font.size = Pt(11)
    quote_font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    heading = doc.add_heading('The Impact of Artificial Intelligence on Global Industries', level=1)
    heading.paragraph_format.space_after = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run('A Comprehensive Analysis of AI Adoption Across Key Economic Sectors')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(INTRO_TEXT)

    # --- Body sections with quotes interspersed ---
    section_titles = [
        '2. Healthcare and Biomedical Research',
        '3. Financial Services and Regulatory Compliance',
        '4. Manufacturing and Quality Assurance',
        '5. Education Technology and Adaptive Learning',
        '6. Agriculture and Precision Farming',
        '7. Retail and Consumer Personalization',
        '8. Energy and Grid Optimization',
        '9. Legal Services and Document Analysis',
        '10. Transportation and Logistics',
        '11. Cybersecurity and Threat Detection',
    ]

    quote_idx = 0
    for i, (title, body_text) in enumerate(zip(section_titles, BODY_PARAGRAPHS)):
        doc.add_heading(title, level=2)
        doc.add_paragraph(body_text)

        # Add a repeated/expanded version to increase page count
        doc.add_paragraph(
            f"Further analysis of the {title.split('. ', 1)[1].lower()} sector reveals "
            f"additional trends that merit discussion. Industry analysts predict continued "
            f"growth in AI adoption rates through 2027, with compound annual growth rates "
            f"exceeding 25% in most sub-segments. The competitive pressure to adopt these "
            f"technologies has created a robust market for AI consulting services, with "
            f"firms like Accenture, Deloitte, and Boston Consulting Group each reporting "
            f"over $1 billion in AI-related consulting revenue for the 2024 fiscal year."
        )

        # Insert a block quote every 2 sections
        if i % 2 == 1 and quote_idx < len(QUOTE_BLOCKS):
            p = doc.add_paragraph(QUOTE_BLOCKS[quote_idx], style='Quotations')
            quote_idx += 1

    # Insert remaining quotes
    while quote_idx < len(QUOTE_BLOCKS):
        doc.add_heading(f'{12 + quote_idx - len(QUOTE_BLOCKS) + len(QUOTE_BLOCKS)}. Expert Perspectives', level=2)
        doc.add_paragraph(
            "Leading experts in the field have offered their perspectives on the "
            "current state and future trajectory of artificial intelligence technologies."
        )
        p = doc.add_paragraph(QUOTE_BLOCKS[quote_idx], style='Quotations')
        quote_idx += 1

    # --- Closing sections ---
    doc.add_heading('12. Future Outlook and Emerging Trends', level=2)
    for cp in CLOSING_PARAGRAPHS:
        doc.add_paragraph(cp)

    # Add additional filler content to reach ~30 pages
    additional_topics = [
        ('13. Government and Public Administration',
         'Government agencies at federal, state, and municipal levels have begun deploying AI '
         'systems for a wide range of administrative functions, from processing tax returns to '
         'optimizing public transportation routes. The U.S. Internal Revenue Service piloted an '
         'AI-assisted audit selection system in 2024 that improved detection of fraudulent '
         'returns by 45% while reducing the number of audits conducted on compliant taxpayers '
         'by 30%. Similar systems have been deployed by tax authorities in the United Kingdom, '
         'Australia, and Singapore, with comparable improvements in efficiency and accuracy.'),
        ('14. Entertainment and Creative Industries',
         'The entertainment industry has experienced a profound transformation as generative AI '
         'tools have become integral to content production workflows. Major studios including '
         'Warner Bros., Disney, and Netflix have invested over $3 billion collectively in AI '
         'research and development since 2023. These investments have yielded tools that assist '
         'with script analysis, visual effects generation, music composition, and audience '
         'engagement prediction. The Screen Actors Guild reported that AI-generated background '
         'characters now appear in 40% of major film productions, reducing production costs '
         'while raising complex questions about creative ownership and artistic authenticity.'),
        ('15. Real Estate and Urban Planning',
         'Urban planners and real estate developers have adopted AI-powered tools for site '
         'selection, demand forecasting, and building design optimization. Zillow Group\'s '
         'Zestimate algorithm, enhanced with transformer-based architecture in 2024, now '
         'achieves a median error rate of just 1.9% for on-market properties, making it the '
         'most accurate publicly available home valuation tool. Commercial real estate firms '
         'such as CBRE and JLL use predictive models to identify emerging market trends, '
         'optimize lease negotiations, and forecast occupancy rates for office and retail spaces.'),
        ('16. Telecommunications and Network Infrastructure',
         'Telecommunications providers have deployed AI systems for network optimization, '
         'customer service automation, and 5G infrastructure planning. AT&T reported that its '
         'AI-driven network management platform reduced service outages by 35% in 2024 while '
         'improving average data throughput by 22% during peak usage periods. Verizon and '
         'T-Mobile have made similar investments, with combined capital expenditures on AI '
         'infrastructure exceeding $8 billion for the year. The deployment of AI at the edge '
         'of the network has enabled new applications in autonomous vehicles, remote surgery, '
         'and augmented reality that require ultra-low latency connections.'),
        ('17. Insurance and Risk Assessment',
         'The insurance industry has been transformed by machine learning models that enable '
         'more granular risk assessment and claims processing automation. Progressive Insurance '
         'reported that its AI underwriting models, deployed across all personal auto insurance '
         'lines in 2024, improved loss ratio accuracy by 8 percentage points compared to '
         'traditional actuarial methods. Claims adjudication has also been accelerated, with '
         'Lemonade\'s AI claims bot processing simple claims in under three minutes, compared '
         'to an industry average of five to seven business days for traditional carriers.'),
        ('18. Pharmaceutical Research and Drug Discovery',
         'Pharmaceutical companies have embraced AI-driven drug discovery platforms that can '
         'screen millions of molecular compounds in silico, dramatically reducing the time and '
         'cost of identifying promising drug candidates. Insilico Medicine\'s AI platform '
         'identified a novel treatment for idiopathic pulmonary fibrosis that entered Phase II '
         'clinical trials in 2024, having progressed from target identification to clinical '
         'candidate in just 18 months, compared to the industry average of four to six years. '
         'Pfizer, Roche, and Novartis have each established dedicated AI research divisions '
         'with combined annual budgets exceeding $2 billion.'),
        ('19. Environmental Monitoring and Climate Science',
         'Climate scientists and environmental agencies have leveraged AI models to improve '
         'weather forecasting, monitor deforestation, and predict natural disaster impacts. '
         'Google DeepMind\'s GraphCast model, published in late 2023, demonstrated superior '
         'performance to the European Centre for Medium-Range Weather Forecasts for 10-day '
         'predictions at a fraction of the computational cost. NASA\'s Earth Science Division '
         'has deployed convolutional neural networks to analyze satellite imagery, enabling '
         'real-time tracking of Amazon rainforest deforestation with 95% accuracy at 30-meter '
         'resolution, providing critical data for enforcement of environmental protection laws.'),
        ('20. Conclusion and Recommendations',
         'The evidence presented throughout this analysis demonstrates that artificial '
         'intelligence has moved decisively from an experimental technology to a critical '
         'business infrastructure component across virtually every sector of the global economy. '
         'Organizations that have invested early in AI capabilities are already reaping '
         'significant competitive advantages, while laggards face increasing pressure to catch '
         'up. The key recommendations emerging from this analysis are: first, prioritize AI '
         'investments that align with core business objectives rather than pursuing technology '
         'for its own sake; second, invest heavily in workforce development and change management '
         'to ensure successful adoption; third, establish robust governance frameworks that '
         'address ethical considerations, data privacy, and algorithmic accountability; and '
         'fourth, foster collaborative ecosystems that enable knowledge sharing and reduce '
         'duplication of effort across organizations and sectors.'),
    ]

    for title, text in additional_topics:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
        # Add expanded analysis paragraph
        doc.add_paragraph(
            f"The implications of these developments extend beyond immediate operational "
            f"efficiency gains. As AI systems become more sophisticated and ubiquitous, they "
            f"are reshaping competitive dynamics, labor markets, and regulatory landscapes in "
            f"ways that demand sustained attention from business leaders and policymakers alike. "
            f"The organizations that will thrive in this environment are those that view AI not "
            f"as a one-time implementation project but as an ongoing strategic capability that "
            f"requires continuous investment, experimentation, and adaptation to evolving "
            f"technological possibilities and societal expectations."
        )
        # Add a block quote for some sections to use Quotations style
        if title.startswith(('14.', '17.', '19.')):
            doc.add_paragraph(
                "Industry observers have noted that the pace of adoption varies significantly "
                "across geographies and organizational sizes. A comprehensive survey conducted "
                "by Gartner in late 2024 found that 78% of large enterprises had deployed at "
                "least one AI application in production, compared to just 23% of small and "
                "medium-sized businesses. This disparity highlights the need for accessible AI "
                "tools and platforms that democratize access to these transformative capabilities "
                "and prevent the emergence of a widening technology divide.",
                style='Quotations'
            )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
