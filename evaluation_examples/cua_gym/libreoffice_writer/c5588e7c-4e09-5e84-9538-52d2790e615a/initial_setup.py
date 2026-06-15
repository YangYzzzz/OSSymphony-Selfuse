"""
Initial Setup: Run the existing macro 'CleanupFormatting' from the document's macro library.
Task ID: writer_tm_066
Domain: libreoffice_writer

Creates a document with inconsistent spacing, multiple consecutive spaces,
and trailing whitespace. Also installs the CleanupFormatting macro in
LibreOffice's Standard.Module1.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_066'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def install_macro():
    """Install CleanupFormatting macro into LibreOffice Standard.Module1."""
    # LibreOffice user macro path
    macro_dir = os.path.expanduser('~/.config/libreoffice/4/user/basic/Standard')
    os.makedirs(macro_dir, exist_ok=True)

    # Write the Module1.xba macro file
    module_content = '''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE script:module PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "module.dtd">
<script:module xmlns:script="http://openoffice.org/2000/script" script:name="Module1" script:language="StarBasic">
Sub CleanupFormatting()
    Dim oDoc As Object
    Dim oText As Object
    Dim oEnum As Object
    Dim oPar As Object
    Dim oParEnum As Object
    Dim oPortion As Object
    Dim sText As String

    oDoc = ThisComponent
    oText = oDoc.getText()
    oEnum = oText.createEnumeration()

    Do While oEnum.hasMoreElements()
        oPar = oEnum.nextElement()

        If oPar.supportsService("com.sun.star.text.Paragraph") Then
            oParEnum = oPar.createEnumeration()
            Do While oParEnum.hasMoreElements()
                oPortion = oParEnum.nextElement()
                sText = oPortion.getString()
                &apos; Remove trailing whitespace
                Do While Right(sText, 1) = " " Or Right(sText, 1) = Chr(9)
                    sText = Left(sText, Len(sText) - 1)
                Loop
                &apos; Collapse multiple spaces to single
                Do While InStr(sText, "  ") > 0
                    sText = Join(Split(sText, "  "), " ")
                Loop
                oPortion.setString(sText)
            Loop
        End If
    Loop

    MsgBox "Formatting cleanup complete!", 64, "CleanupFormatting"
End Sub
</script:module>'''

    with open(os.path.join(macro_dir, 'Module1.xba'), 'w') as f:
        f.write(module_content)

    # Ensure dialog.xlc and script.xlc exist for the Standard library
    script_xlc = '''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE library:library PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "library.dtd">
<library:library xmlns:library="http://openoffice.org/2000/library" library:name="Standard" library:readonly="false" library:passwordprotected="false">
 <library:element library:name="Module1"/>
</library:library>'''

    with open(os.path.join(macro_dir, 'script.xlc'), 'w') as f:
        f.write(script_xlc)

    dialog_xlc = '''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE library:library PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "library.dtd">
<library:library xmlns:library="http://openoffice.org/2000/library" library:name="Standard" library:readonly="false" library:passwordprotected="false">
</library:library>'''

    with open(os.path.join(macro_dir, 'dialog.xlc'), 'w') as f:
        f.write(dialog_xlc)

    print(f'Macro installed at: {macro_dir}/Module1.xba')


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style.font.size = Pt(12)

    # Title with trailing whitespace
    heading = doc.add_heading('Quarterly  Sales  Report   ', level=1)

    # Paragraph 1 - multiple consecutive spaces and trailing whitespace
    p1 = doc.add_paragraph()
    p1.add_run(
        'The  quarterly  sales  report  for  Q3  2025  shows  significant  '
        'growth  across  all  departments.   Revenue  increased  by  15%  '
        'compared  to  the  previous  quarter,  driven  primarily  by  '
        'strong  performance  in  the  technology  and  healthcare  sectors.   '
    )

    # Paragraph 2 - inconsistent spacing
    p2 = doc.add_paragraph()
    p2.add_run(
        'Key  highlights  from  this  quarter  include:   '
    )

    # Bullet points with messy spacing
    bullets = [
        'Total  revenue  reached  $4.2  million,   up  from  $3.65  million  in  Q2   ',
        'New  client  acquisitions   increased  by  23%  over  the  prior  period   ',
        'Customer  retention  rate   improved  to  94.7%,   the  highest  in  company  history   ',
        'Operating  expenses  were  reduced   by  8%  through  process  optimization   ',
        'Employee  satisfaction   scores  reached  an  all-time  high  of  4.6  out  of  5   ',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Paragraph 3 - department breakdown with extra spaces
    doc.add_heading('Department   Performance   Summary   ', level=2)

    p3 = doc.add_paragraph()
    p3.add_run(
        'The  engineering  department,   led  by  Sarah  Chen,   delivered  '
        'three  major  product  releases   ahead  of  schedule.   The  marketing  '
        'team,   under  Marcus  Johnson,   launched  a  successful  rebranding  '
        'campaign  that  increased  brand  awareness   by  40%.   '
    )

    p4 = doc.add_paragraph()
    p4.add_run(
        'The  finance  department   completed  the  annual  audit   with  no  '
        'material  findings.   David  Park   presented  the  budget  forecast  '
        'for  FY2026,   projecting  a  12%  increase   in  overall  revenue.   '
    )

    # Paragraph 5 - regional data with trailing spaces
    doc.add_heading('Regional   Sales   Data   ', level=2)

    p5 = doc.add_paragraph()
    p5.add_run(
        'North  America   accounted  for  58%  of  total  revenue,   '
        'followed  by  Europe   at  27%  and  Asia-Pacific   at  15%.   '
        'The  APAC  region   showed  the  highest  growth  rate   at  34%  '
        'year-over-year,   suggesting  strong  expansion   potential  in  '
        'emerging  markets.   '
    )

    # Paragraph 6 - outlook
    doc.add_heading('Outlook  and  Recommendations   ', level=2)

    p6 = doc.add_paragraph()
    p6.add_run(
        'Based  on  current  trends   and  market  conditions,   we  '
        'recommend  the  following   strategic  initiatives  for  Q4  2025:   '
    )

    # Numbered list with messy spacing
    items = [
        'Increase  investment   in  the  Asia-Pacific  region   by  20%   ',
        'Launch  the  new   enterprise  product  tier   by  November  15th   ',
        'Hire  15  additional   sales  representatives   for  the  European  market   ',
        'Implement  the  revised   customer  onboarding  process   across  all  regions   ',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Number')

    # Final paragraph
    p7 = doc.add_paragraph()
    p7.add_run(
        'This  report  was  prepared   by  the  Business  Analytics  team.   '
        'For  questions   or  additional  analysis,   please  contact  '
        'analytics@acmecorp.com.   '
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')


# Execute
install_macro()
create_initial()

# Open in LibreOffice Writer
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
