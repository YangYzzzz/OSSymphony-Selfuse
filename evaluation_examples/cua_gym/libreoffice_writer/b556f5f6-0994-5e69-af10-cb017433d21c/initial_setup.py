"""
Initial Setup: Create code_documentation.docx with trailing whitespace on many lines.
Task ID: writer_edit_015
Domain: libreoffice_writer

The document is a 3-page technical code documentation file.
Many paragraphs/lines have 1-4 trailing spaces (as if copy-pasted from a code editor).
The agent's task is to use regex Find & Replace to remove all trailing whitespace.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'code_documentation'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # ---- Page 1: Module Overview ----

    # Title heading (no trailing space — title is clean)
    h = doc.add_heading('DataProcessor Module — API Reference', level=1)

    # Intro paragraph with trailing spaces on some runs
    p = doc.add_paragraph()
    run = p.add_run('This module provides utilities for processing, transforming, and exporting   ')
    run2 = p.add_run('structured data from CSV files and database queries.')
    # run2 intentionally no trailing space

    p2 = doc.add_paragraph()
    run2a = p2.add_run('Version: 2.4.1  ')
    run2b = p2.add_run('  Author: Elena Vasquez   ')
    run2c = p2.add_run('License: MIT')

    # Section heading
    doc.add_heading('Installation', level=2)

    p3 = doc.add_paragraph()
    p3.add_run('Install via pip: ')
    run3b = p3.add_run('pip install dataprocessor   ')

    p4 = doc.add_paragraph()
    p4.add_run('Requires Python 3.8 or higher.  ')

    p5 = doc.add_paragraph()
    p5.add_run('Dependencies:  ')

    # Bullet-style entries
    pb1 = doc.add_paragraph(style='List Bullet')
    pb1.add_run('pandas >= 1.3.0   ')

    pb2 = doc.add_paragraph(style='List Bullet')
    pb2.add_run('numpy >= 1.21.0 ')

    pb3 = doc.add_paragraph(style='List Bullet')
    pb3.add_run('sqlalchemy >= 1.4  ')

    pb4 = doc.add_paragraph(style='List Bullet')
    pb4.add_run('pydantic >= 1.8.2')

    # Section heading
    doc.add_heading('Quick Start', level=2)

    p6 = doc.add_paragraph()
    p6.add_run('The following example demonstrates basic usage of the DataProcessor class:  ')

    # Code block (indented paragraph)
    code_lines = [
        'from dataprocessor import DataProcessor   ',
        '',
        'dp = DataProcessor(source="sales_data.csv")  ',
        'dp.load()   ',
        'dp.clean(drop_nulls=True, strip_whitespace=True)  ',
        'result = dp.export(format="json")   ',
        'print(result.summary())  ',
    ]
    for line in code_lines:
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Pt(36)
        run_code = cp.add_run(line)
        run_code.font.name = 'Courier New'
        run_code.font.size = Pt(10)

    # Page break
    doc.add_page_break()

    # ---- Page 2: Class Reference ----

    doc.add_heading('Class Reference', level=1)

    doc.add_heading('DataProcessor', level=2)

    p7 = doc.add_paragraph()
    p7.add_run('class ')
    r7b = p7.add_run('DataProcessor')
    r7b.bold = True
    r7b.font.name = 'Courier New'
    p7.add_run('(source, config=None, verbose=False)  ')

    p8 = doc.add_paragraph()
    p8.add_run('Initializes a new DataProcessor instance for the given data source.  ')

    doc.add_heading('Parameters', level=3)

    params = [
        ('source', 'str', 'Path to a CSV file or a database URI string.   '),
        ('config', 'dict, optional', 'Configuration dictionary for processing options. Defaults to None.  '),
        ('verbose', 'bool, optional', 'Enable verbose logging output. Defaults to False.   '),
    ]
    for name, ptype, desc in params:
        pp = doc.add_paragraph(style='List Bullet')
        rn = pp.add_run(name)
        rn.bold = True
        rn.font.name = 'Courier New'
        pp.add_run(f' ({ptype}): {desc}')

    doc.add_heading('Methods', level=2)

    methods = [
        ('load()', 'Load data from the specified source into memory.   '),
        ('clean(**kwargs)', 'Apply a series of data cleaning operations.  '),
        ('transform(pipeline)', 'Run a custom transformation pipeline on the data.   '),
        ('validate(schema)', 'Validate data against a Pydantic schema model.  '),
        ('export(format, path=None)', 'Export processed data to the desired output format.   '),
        ('summary()', 'Return a statistical summary of the loaded data.  '),
        ('reset()', 'Reset the processor to its initial state.   '),
    ]

    for method_sig, method_desc in methods:
        mp = doc.add_paragraph()
        mr1 = mp.add_run(method_sig)
        mr1.bold = True
        mr1.font.name = 'Courier New'
        mr1.font.size = Pt(11)
        mp2 = doc.add_paragraph()
        mp2.paragraph_format.left_indent = Pt(18)
        mp2.add_run(method_desc)

    # Page break
    doc.add_page_break()

    # ---- Page 3: Configuration & Examples ----

    doc.add_heading('Configuration Reference', level=1)

    p9 = doc.add_paragraph()
    p9.add_run('The config dictionary accepts the following keys:  ')

    config_items = [
        ('encoding', '"utf-8"', 'File encoding for CSV source files.   '),
        ('delimiter', '","', 'Column delimiter character.  '),
        ('header_row', '0', 'Zero-based row index of the header.   '),
        ('null_values', '["", "N/A", "NULL"]', 'List of strings treated as null.  '),
        ('date_columns', '[]', 'Column names to parse as datetime objects.   '),
        ('chunk_size', '10000', 'Number of rows per processing chunk.  '),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Key'
    hdr_cells[1].text = 'Default'
    hdr_cells[2].text = 'Description'
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for key, default, description in config_items:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = default
        row_cells[2].text = description

    doc.add_heading('Advanced Example', level=2)

    p10 = doc.add_paragraph()
    p10.add_run('The following example demonstrates error handling and validation:   ')

    advanced_code = [
        'from dataprocessor import DataProcessor, ProcessingError  ',
        'from myschemas import SalesRecord  ',
        '',
        'config = {  ',
        '    "encoding": "utf-8",  ',
        '    "null_values": ["", "N/A", "—"],   ',
        '    "date_columns": ["order_date", "ship_date"],  ',
        '}  ',
        '',
        'try:  ',
        '    dp = DataProcessor("quarterly_sales.csv", config=config)  ',
        '    dp.load()   ',
        '    dp.clean(drop_nulls=True, remove_duplicates=True)  ',
        '    dp.validate(SalesRecord)  ',
        '    output = dp.export("parquet", path="/data/output/q2_sales.parquet")  ',
        '    print(f"Exported {output.row_count} rows successfully.")   ',
        'except ProcessingError as e:   ',
        '    print(f"Processing failed: {e.message}")  ',
        '    raise  ',
    ]

    for line in advanced_code:
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Pt(36)
        run_code = cp.add_run(line)
        run_code.font.name = 'Courier New'
        run_code.font.size = Pt(10)

    doc.add_heading('Changelog', level=2)

    changelog_entries = [
        'v2.4.1 — Fixed edge case in null detection for mixed-type columns.   ',
        'v2.4.0 — Added support for Parquet and Feather export formats.  ',
        'v2.3.2 — Performance improvements for large CSV files (>1M rows).   ',
        'v2.3.0 — Introduced pipeline-based transform API.  ',
        'v2.2.1 — Bug fix: date parsing with missing timezone info.   ',
        'v2.2.0 — Added Pydantic validation support.  ',
    ]

    for entry in changelog_entries:
        cp = doc.add_paragraph(style='List Bullet')
        cp.add_run(entry)

    p_end = doc.add_paragraph()
    p_end.add_run('For full changelog and migration guide, visit the project repository.  ')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
