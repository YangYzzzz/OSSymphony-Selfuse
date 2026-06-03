"""
Reward Script: SmartHome Hub User Manual creation
Task ID: writer_wf_008
Domain: libreoffice_writer
Scoring:
  Component 1: Title page with product name and version 2.0 (0.15)
  Component 2: Table of Contents section (0.10)
  Component 3: Chapter 1 Getting Started as Heading 1 (0.10)
  Component 4: Chapter 2 Installation as Heading 1 (0.10)
  Component 5: Chapter 3 Configuration as Heading 1 (0.10)
  Component 6: Chapter 4 Troubleshooting as Heading 1 (0.10)
  Component 7: Troubleshooting table with 5 issue rows (0.20)
  Component 8: WARNING box with red border and safety notice (0.15)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_008'


def persist_app_state(domain):
    """Try to save any unsaved changes in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts and styles for reuse
    all_paras = [(p.text.strip(), p.style.name if p.style else 'None') for p in doc.paragraphs]
    h1_texts = [text for text, style in all_paras if style == 'Heading 1']

    # Component 1: Title page with "SmartHome Hub" and "Version 2.0" (0.15 points)
    try:
        # Check that somewhere in the document text there is "SmartHome Hub" and version 2.0
        all_text_lower = ' '.join(t for t, s in all_paras).lower()
        has_product_name = 'smarthome hub' in all_text_lower
        has_version = '2.0' in all_text_lower
        if has_product_name and has_version:
            print(f"PASS: Component 1 — Title page has 'SmartHome Hub' and 'Version 2.0' (0.15 pts)")
            total_score += 0.15
        elif has_product_name:
            print(f"PARTIAL: Component 1 — Has product name but missing version 2.0 (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 1 — Missing product name 'SmartHome Hub' (has_name={has_product_name}, has_ver={has_version})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table of Contents section (0.10 points)
    try:
        toc_found = any('table of contents' in text.lower() for text, style in all_paras)
        if toc_found:
            print(f"PASS: Component 2 — Table of Contents section found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — No 'Table of Contents' text found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Chapter 1: Getting Started as Heading 1 (0.10 points)
    try:
        ch1_found = any('getting started' in h.lower() for h in h1_texts)
        if ch1_found:
            print(f"PASS: Component 3 — Chapter 1 'Getting Started' as Heading 1 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No Heading 1 containing 'Getting Started'. H1s found: {h1_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Chapter 2: Installation as Heading 1 (0.10 points)
    try:
        ch2_found = any('installation' in h.lower() for h in h1_texts)
        if ch2_found:
            print(f"PASS: Component 4 — Chapter 2 'Installation' as Heading 1 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — No Heading 1 containing 'Installation'. H1s found: {h1_texts}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Chapter 3: Configuration as Heading 1 (0.10 points)
    try:
        ch3_found = any('configuration' in h.lower() for h in h1_texts)
        if ch3_found:
            print(f"PASS: Component 5 — Chapter 3 'Configuration' as Heading 1 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — No Heading 1 containing 'Configuration'. H1s found: {h1_texts}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Chapter 4: Troubleshooting as Heading 1 (0.10 points)
    try:
        ch4_found = any('troubleshooting' in h.lower() for h in h1_texts)
        if ch4_found:
            print(f"PASS: Component 6 — Chapter 4 'Troubleshooting' as Heading 1 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — No Heading 1 containing 'Troubleshooting'. H1s found: {h1_texts}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Troubleshooting table with 5 issue rows (0.20 points)
    # The table should have a header row + 5 data rows = 6 total rows, with 3 columns (Issue, Cause, Solution)
    try:
        found_trouble_table = False
        for table in doc.tables:
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            # Look for a table with at least 6 rows (header + 5 issues) and at least 2 columns
            if num_rows >= 6 and num_cols >= 2:
                # Check if header row contains issue-related headers
                header_text = ' '.join(c.text.strip().lower() for c in table.rows[0].cells)
                if 'issue' in header_text or 'problem' in header_text or 'cause' in header_text or 'solution' in header_text:
                    found_trouble_table = True
                    # Count data rows (non-empty first column after header)
                    data_rows = sum(1 for r in range(1, num_rows) if table.cell(r, 0).text.strip())
                    if data_rows >= 5:
                        print(f"PASS: Component 7 — Troubleshooting table has {data_rows} data rows, {num_cols} cols (0.20 pts)")
                        total_score += 0.20
                    else:
                        partial = round(0.20 * min(data_rows, 5) / 5, 2)
                        print(f"PARTIAL: Component 7 — Found table but only {data_rows}/5 data rows ({partial} pts)")
                        total_score += partial
                    break

        if not found_trouble_table:
            # Fallback: check any table with >= 6 rows
            for table in doc.tables:
                if len(table.rows) >= 6 and len(table.columns) >= 2:
                    found_trouble_table = True
                    print(f"PARTIAL: Component 7 — Found a table with {len(table.rows)} rows but unclear headers (0.10 pts)")
                    total_score += 0.10
                    break

        if not found_trouble_table:
            print(f"FAIL: Component 7 — No troubleshooting table found. Tables in doc: {len(doc.tables)}")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: WARNING box with red border and safety notice (0.15 points)
    # In the golden file, this is implemented as a single-cell table with red (FF0000) borders
    try:
        warning_found = False
        for table in doc.tables:
            if len(table.rows) == 1 and len(table.columns) == 1:
                cell = table.cell(0, 0)
                cell_text = cell.text.strip().lower()
                if 'warning' in cell_text or 'safety' in cell_text:
                    # Check for red border
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    has_red_border = False
                    if tcPr is not None:
                        borders = tcPr.find(qn('w:tcBorders'))
                        if borders is not None:
                            for border in borders:
                                color = border.get(qn('w:color'))
                                if color and color.upper() in ('FF0000', 'RED', 'CC0000', 'DD0000', 'EE0000', 'BB0000'):
                                    has_red_border = True
                                    break

                    # Also check table-level borders
                    if not has_red_border:
                        tbl = table._element
                        tblPr = tbl.find(qn('w:tblPr'))
                        if tblPr is not None:
                            tbl_borders = tblPr.find(qn('w:tblBorders'))
                            if tbl_borders is not None:
                                for border in tbl_borders:
                                    color = border.get(qn('w:color'))
                                    if color and color.upper() in ('FF0000', 'RED', 'CC0000', 'DD0000', 'EE0000', 'BB0000'):
                                        has_red_border = True
                                        break

                    if has_red_border:
                        print(f"PASS: Component 8 — WARNING box with red border and safety notice (0.15 pts)")
                        total_score += 0.15
                        warning_found = True
                    else:
                        # Has warning text but no red border
                        print(f"PARTIAL: Component 8 — WARNING text found but no red border detected (0.07 pts)")
                        total_score += 0.07
                        warning_found = True
                    break

        if not warning_found:
            # Broader check: any table or paragraph with WARNING
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if 'warning' in cell.text.lower() and 'safety' in cell.text.lower():
                            print(f"PARTIAL: Component 8 — Found WARNING/safety text in a table but not in expected format (0.05 pts)")
                            total_score += 0.05
                            warning_found = True
                            break
                    if warning_found:
                        break
                if warning_found:
                    break

        if not warning_found:
            print(f"FAIL: Component 8 — No WARNING box with safety notice found")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
