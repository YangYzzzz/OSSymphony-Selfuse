"""
FINAL REWARD SCRIPT - SUCCESS
Task: Please center the heading that appears directly under the logo on the cover page.
Generated: 2025-10-14 11:22:35
Status: success
Model: azure-o3
Total Steps: 2
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def verify_center_heading(file_path):
    """Reward script to verify that the first text heading directly under the logo
    on the cover page is centered.
    Progressive scoring:
      • 0.3 points for successfully locating the candidate heading paragraph.
      • 0.7 additional points if that heading is center-aligned.
    Returns a float between 0.0 and 1.0 and prints detailed diagnostics.
    """

    print(f"Starting verification for: {file_path}")
    max_score = 1.0
    score = 0.0

    # ---------- Prerequisite checks (NO POINTS) ----------
    if not os.path.exists(file_path):
        print("✗ File does not exist.")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
        print(f"✓ Document loaded successfully with {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Failed to load DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- Locate the heading under the logo ----------
    # Heuristic: the first non-empty paragraph (images have no paragraph text)
    heading_para = None
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            heading_para = para
            break

    if heading_para is None:
        print("✗ No non-empty text paragraph found – cannot verify heading.")
        print("REWARD: 0.0")
        return 0.0

    print(f"✓ Candidate heading found: '{heading_para.text.strip()[:60]}'")
    score += 0.3  # credit for locating the heading

    # ---------- Verify alignment ----------
    alignment = heading_para.paragraph_format.alignment  # may be None if not set
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        print("✓ Heading is centered")
        score += 0.7
    else:
        print(f"✗ Heading is not centered (alignment={alignment})")

    # ---------- Final score ----------
    final_score = min(score, max_score)
    print(f"Total score: {final_score}/{max_score}")
    print(f"REWARD: {final_score}")
    return final_score

if __name__ == "__main__":
    file_path = "/home/user/please_center_the_heading_that_appears_directly_under_the_logo_on_the_cover_page.docx"
    verify_center_heading(file_path)
