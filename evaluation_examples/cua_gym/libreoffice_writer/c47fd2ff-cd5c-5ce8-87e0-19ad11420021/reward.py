"""
Reward Script: Monthly Sales Report in LibreOffice Writer
Task ID: writer_wf_090
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Title "Monthly Sales Report - September 2025"
  Component 2 (0.10): Report info (prepared by, department, date)
  Component 3 (0.15): Six Heading 1 sections with correct names
  Component 4 (0.20): Regional sales table (6 rows x 5 cols)
  Component 5 (0.15): Top products table (6 rows x 3 cols)
  Component 6 (0.10): Customer acquisition stats (new + returning)
  Component 7 (0.15): Logo placeholder in header
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_090'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title paragraph (0.15 points)
    # The document must have a Title-styled paragraph containing the report title.
    try:
        title_found = False
        for para in doc.paragraphs:
            if para.style and para.style.name == 'Title':
                if 'monthly sales report' in para.text.lower() and 'september 2025' in para.text.lower():
                    title_found = True
                    break
        if title_found:
            print(f"PASS: Component 1 — Title paragraph found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — No Title-styled paragraph with 'Monthly Sales Report - September 2025'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Report info lines — prepared by, department, date (0.10 points)
    # Must have paragraphs containing these three pieces of info.
    try:
        all_text = ' '.join(p.text.lower() for p in doc.paragraphs)
        has_prepared = 'prepared by' in all_text or 'prepared:' in all_text
        has_department = 'department' in all_text
        has_date = 'date' in all_text or 'october' in all_text or '2025' in all_text
        info_count = sum([has_prepared, has_department, has_date])
        if info_count == 3:
            print(f"PASS: Component 2 — All 3 report info lines found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Only {info_count}/3 info lines (prepared={has_prepared}, dept={has_department}, date={has_date})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Six Heading 1 sections (0.15 points)
    # Required sections: Executive Summary, Sales by Region, Top Products,
    # Customer Acquisition, Challenges, Outlook and Targets for Next Month
    try:
        required_sections = [
            'executive summary',
            'sales by region',
            'top products',
            'customer acquisition',
            'challenges',
            'outlook',
        ]
        heading1_texts = []
        for para in doc.paragraphs:
            if para.style and para.style.name == 'Heading 1':
                heading1_texts.append(para.text.lower().strip())

        matched = 0
        for req in required_sections:
            if any(req in h for h in heading1_texts):
                matched += 1

        if matched >= 6:
            print(f"PASS: Component 3 — All 6 Heading 1 sections found (0.15 pts)")
            total_score += 0.15
        elif matched >= 4:
            partial = round(0.15 * matched / 6, 3)
            print(f"PARTIAL: Component 3 — {matched}/6 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {matched}/6 Heading 1 sections found: {heading1_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Regional sales table (0.20 points)
    # Must have a table with 6 rows (header + 4 regions + total) and 5 columns
    # Columns: Region, Target, Actual, Variance, % Achievement
    try:
        region_table_found = False
        for table in doc.tables:
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows < 5 or num_cols < 5:
                continue
            # Check if header row matches expected columns
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            has_region = any('region' in c for c in header_cells)
            has_target = any('target' in c for c in header_cells)
            has_actual = any('actual' in c for c in header_cells)
            if has_region and has_target and has_actual:
                # Check for Total row
                last_row_cells = [cell.text.strip().lower() for cell in table.rows[-1].cells]
                has_total = any('total' in c for c in last_row_cells)
                # Check that there are at least 4 region data rows
                data_rows = num_rows - 1  # minus header
                if has_total:
                    data_rows -= 1  # minus total row
                if data_rows >= 4 and has_total:
                    region_table_found = True
                    print(f"PASS: Component 4 — Regional sales table: {num_rows} rows x {num_cols} cols with Total row (0.20 pts)")
                    total_score += 0.20
                elif data_rows >= 4:
                    # Has regions but no total
                    print(f"PARTIAL: Component 4 — Regional table found but missing Total row (0.12 pts)")
                    total_score += 0.12
                    region_table_found = True
                elif has_total:
                    # Has total but fewer regions
                    print(f"PARTIAL: Component 4 — Regional table found but only {data_rows} regions (0.12 pts)")
                    total_score += 0.12
                    region_table_found = True
                break
        if not region_table_found:
            print(f"FAIL: Component 4 — No regional sales table with 5+ cols and Region/Target/Actual headers")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Top products table (0.15 points)
    # Must have a table with header + 5 product rows, columns: Product, Units Sold, Revenue
    try:
        products_table_found = False
        for table in doc.tables:
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows < 4 or num_cols < 3:
                continue
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            has_product = any('product' in c for c in header_cells)
            has_units = any('unit' in c for c in header_cells)
            has_revenue = any('revenue' in c for c in header_cells)
            if has_product and has_units and has_revenue:
                # Count data rows (non-header, non-empty)
                data_count = 0
                for ri in range(1, num_rows):
                    cell_text = table.rows[ri].cells[0].text.strip()
                    if cell_text:
                        data_count += 1
                if data_count >= 5:
                    print(f"PASS: Component 5 — Top products table: {num_rows} rows x {num_cols} cols, {data_count} products (0.15 pts)")
                    total_score += 0.15
                    products_table_found = True
                elif data_count >= 3:
                    partial = round(0.15 * data_count / 5, 3)
                    print(f"PARTIAL: Component 5 — Products table found with {data_count}/5 products ({partial} pts)")
                    total_score += partial
                    products_table_found = True
                break
        if not products_table_found:
            print(f"FAIL: Component 5 — No products table with Product/Units/Revenue headers found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Customer acquisition stats (0.10 points)
    # Must mention new customers and returning customers in the document text
    try:
        all_text_lower = ' '.join(p.text.lower() for p in doc.paragraphs)
        has_new_customers = 'new customer' in all_text_lower
        has_returning = 'returning customer' in all_text_lower or 'retention' in all_text_lower
        if has_new_customers and has_returning:
            print(f"PASS: Component 6 — Customer acquisition stats (new + returning) found (0.10 pts)")
            total_score += 0.10
        elif has_new_customers or has_returning:
            print(f"PARTIAL: Component 6 — Only partial customer stats (new={has_new_customers}, returning={has_returning}) (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No customer acquisition stats found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Logo placeholder in header (0.15 points)
    # The document header must contain a logo placeholder (text like "[Company Logo]" or an image)
    try:
        header_has_logo = False
        for section in doc.sections:
            header = section.header
            if header.is_linked_to_previous:
                continue
            for para in header.paragraphs:
                # Check for text placeholder
                if 'logo' in para.text.lower() or 'company' in para.text.lower():
                    header_has_logo = True
                    break
                # Check for image in header
                for run in para.runs:
                    if 'graphicData' in run._element.xml or 'blip' in run._element.xml:
                        header_has_logo = True
                        break
            if header_has_logo:
                break

        if header_has_logo:
            print(f"PASS: Component 7 — Logo placeholder found in header (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 7 — No logo placeholder in document header")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Run persistence before verification
persist_app_state()

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
