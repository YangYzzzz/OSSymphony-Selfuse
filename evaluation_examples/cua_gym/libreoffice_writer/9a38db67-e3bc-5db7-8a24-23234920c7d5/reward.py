"""
FINAL REWARD SCRIPT - SUCCESS
Task: In LibreOffice Writer I’ve got a table labeled “Table 1.” For its third column I want every data cell right-aligned, but I still need the header cell for that same column sitting dead-center. What’s the quickest way to do that?
Generated: 2025-09-10 13:17:02
Status: success
Model: azure-o3
Total Steps: 8
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re


def find_table_by_caption(doc, caption_pattern: str = r"^table\s*1\b"):
    """Return the first Table object whose preceding paragraph matches the caption pattern.
    Fallback: return None if not found (caller can decide what to do)."""
    caption_re = re.compile(caption_pattern, re.IGNORECASE)
    body = doc.element.body
    prev_para_text = None
    tbl_index = 0  # maps XML <w:tbl> order to doc.tables order

    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith('}p'):
            # Collect raw text of this paragraph (possible caption)
            texts = [t.text for t in child.iter() if t.tag.endswith('}t') and t.text]
            prev_para_text = ''.join(texts).strip()
        elif tag.endswith('}tbl'):
            # Encountered a table; check if previous paragraph is desired caption
            match = prev_para_text and caption_re.match(prev_para_text)
            tbl_obj = doc.tables[tbl_index] if tbl_index < len(doc.tables) else None
            if match:
                return tbl_obj
            tbl_index += 1
            prev_para_text = None  # reset after table
    return None


def verify_table_alignment(file_path: str) -> float:
    """Verification logic for the Writer task.

    Requirements:
    1. In the table labelled "Table 1", *header* cell of 3rd column must be CENTER aligned.
    2. *Data* cells (rows below header) in the same column must be RIGHT aligned.

    Scoring (progressive):
        Header centred  -> 0.5 points
        Data cells      -> up to 0.5 points, proportional to how many are right-aligned
    """
    max_score = 1.0
    total_score = 0.0

    print(f"Verifying task for file: {file_path}\n")

    # ---------- Prerequisite checks ----------
    if not os.path.exists(file_path):
        print("✗ File not found – task incomplete")
        return 0.0
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to load DOCX file: {e}")
        return 0.0  # cannot proceed

    # ---------- Locate target table ----------
    target_table = find_table_by_caption(doc)
    if target_table is None:
        print("⚠️ Could not locate table with caption 'Table 1'. Falling back to first table.")
        if doc.tables:
            target_table = doc.tables[0]
        else:
            print("✗ No tables found in the document")
            return 0.0
    else:
        print("✓ Located table labeled 'Table 1'.")

    # Ensure table has at least 3 columns to verify third column
    if len(target_table.columns) < 3:
        print("✗ Target table does not have at least 3 columns – cannot verify")
        return 0.0

    # ---------- Requirement 1: Header centred ----------
    header_cell = target_table.rows[0].cells[2]  # third column, first row
    header_centered = all(
        para.alignment == WD_ALIGN_PARAGRAPH.CENTER for para in header_cell.paragraphs
    )
    if header_centered:
        print("✓ Header cell in third column is CENTER aligned (0.5 points)")
        total_score += 0.5
    else:
        print("✗ Header cell in third column is NOT center-aligned (0 points)")

    # ---------- Requirement 2: Data cells right-aligned ----------
    data_cells = 0
    right_aligned_cells = 0
    for row in target_table.rows[1:]:  # skip header row
        cell = row.cells[2]
        data_cells += 1
        if all(para.alignment == WD_ALIGN_PARAGRAPH.RIGHT for para in cell.paragraphs):
            right_aligned_cells += 1

    if data_cells == 0:
        print("✗ No data rows found in the table – cannot verify data alignment")
    else:
        proportion = right_aligned_cells / data_cells
        data_score = 0.5 * proportion  # proportional credit
        total_score += data_score
        print(f"✓ {right_aligned_cells}/{data_cells} data cells right-aligned (+{data_score:.2f} points)")

    # ---------- Final score ----------
    final_score = min(total_score, max_score)
    print(f"\nTotal Score: {final_score} (out of {max_score})")
    return final_score


if __name__ == "__main__":
    DOC_PATH = "/home/user/in_libreoffice_writer_ive_got_a_table_labeled_table_1_for_its_third_column_i_want_every_data_cell_ri.docx"
    reward = verify_table_alignment(DOC_PATH)
    print(f"REWARD: {reward}")
