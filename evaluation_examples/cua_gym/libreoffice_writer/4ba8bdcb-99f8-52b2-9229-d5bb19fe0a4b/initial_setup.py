"""
Initial Setup: Linear algebra document with placeholder for identity matrix equation
Task ID: writer_acad_082
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_082'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading("Introduction to Linear Algebra", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author line ---
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("Dr. Elena Vasquez")
    run.font.size = Pt(11)
    run.font.italic = True
    run2 = author.add_run("\nDepartment of Mathematics, Westfield University")
    run2.font.size = Pt(10)
    run2.font.italic = True

    # --- Blank line ---
    doc.add_paragraph()

    # --- Section 1: Matrices ---
    doc.add_heading("1. Matrices and Their Properties", level=2)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run(
        "A matrix is a rectangular array of numbers, symbols, or expressions "
        "arranged in rows and columns. Matrices are fundamental objects in linear "
        "algebra and have widespread applications in physics, engineering, computer "
        "science, and economics. The study of matrices encompasses operations such "
        "as addition, scalar multiplication, transposition, and matrix multiplication."
    )
    r1.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(
        "Matrices can be classified by their dimensions (m x n), where m denotes the "
        "number of rows and n the number of columns. A square matrix has equal numbers "
        "of rows and columns. Special types of square matrices include diagonal matrices, "
        "symmetric matrices, and the identity matrix, each possessing unique algebraic "
        "properties that simplify computations in various contexts."
    )
    r2.font.size = Pt(11)

    # --- Section 2: Identity Matrix ---
    doc.add_heading("2. The Identity Matrix", level=2)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run(
        "The identity matrix, commonly denoted as I or I_n, is a square matrix with "
        "ones on the main diagonal and zeros in all other positions. It serves as the "
        "multiplicative identity in matrix algebra: for any compatible matrix A, the "
        "products AI = IA = A hold. The identity matrix plays a central role in the "
        "theory of invertible matrices and linear transformations."
    )
    r3.font.size = Pt(11)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    r4 = p4.add_run(
        "Below is the standard representation of the 3x3 identity matrix:"
    )
    r4.font.size = Pt(11)

    # --- Placeholder for equation ---
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    placeholder.paragraph_format.space_before = Pt(12)
    placeholder.paragraph_format.space_after = Pt(12)
    pr = placeholder.add_run("[Insert 3x3 Identity Matrix Equation Here]")
    pr.font.size = Pt(11)
    pr.font.italic = True
    pr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # --- Continuation ---
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(6)
    r5 = p5.add_run(
        "The 3x3 identity matrix is particularly significant in three-dimensional "
        "geometry and physics, where it represents the transformation that leaves "
        "every vector unchanged. When applied to coordinate transformations, rotations "
        "are often expressed as orthogonal matrices whose product with their transpose "
        "yields the identity matrix."
    )
    r5.font.size = Pt(11)

    # --- Section 3: Applications ---
    doc.add_heading("3. Applications in Computational Mathematics", level=2)

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(6)
    r6 = p6.add_run(
        "Identity matrices are extensively used in numerical algorithms for solving "
        "systems of linear equations. Gaussian elimination, LU decomposition, and "
        "iterative methods such as Jacobi and Gauss-Seidel iterations all rely on "
        "the properties of the identity matrix. In computer graphics, the identity "
        "matrix serves as the default transformation matrix, representing no change "
        "to the object's position, rotation, or scale."
    )
    r6.font.size = Pt(11)

    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(6)
    r7 = p7.add_run(
        "Furthermore, the eigenvalue problem Ax = \u03bbx can be reformulated as "
        "(A - \u03bbI)x = 0, demonstrating the identity matrix's role in spectral "
        "analysis. The characteristic polynomial det(A - \u03bbI) = 0 determines the "
        "eigenvalues of a matrix, which encode fundamental information about linear "
        "transformations including stability, oscillation modes, and principal components."
    )
    r7.font.size = Pt(11)

    # --- References ---
    doc.add_heading("References", level=2)

    refs = [
        "Strang, G. (2016). Introduction to Linear Algebra (5th ed.). Wellesley-Cambridge Press.",
        "Lay, D. C., Lay, S. R., & McDonald, J. J. (2021). Linear Algebra and Its Applications (6th ed.). Pearson.",
        "Horn, R. A., & Johnson, C. R. (2012). Matrix Analysis (2nd ed.). Cambridge University Press.",
    ]
    for i, ref in enumerate(refs, 1):
        rp = doc.add_paragraph()
        rr = rp.add_run(f"[{i}] {ref}")
        rr.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
