"""
FINAL REWARD SCRIPT - SUCCESS
Task: I just dumped a 12-page interview into LibreOffice Writer and every single quotation mark is a plain straight ASCII character—double quote U+0022 and single quote U+0027. Is there a one-step way to turn all of those into proper curly quotes for U.S. English, i.e. “ ” for doubles and ‘ ’ for singles, without running a million find-and-replaces?
Generated: 2025-09-10 16:38:47
Status: success
Model: azure-o3
Total Steps: 2
"""

import os
from docx import Document

# ------------------------------------------------------------
# Reward Script: Verify Curly Quote Replacement in DOCX
# ------------------------------------------------------------
# Task Context:
# The user had a document containing ONLY straight ASCII quotes (" and ').
# The task was to convert them—ideally in one step—into the proper
# U.S. English curly quotes:
#   • “ and ”  (U+201C, U+201D)  for double quotes
#   • ‘ and ’  (U+2018, U+2019)  for single quotes
# ------------------------------------------------------------
# Verification Logic & Progressive Scoring (0.0 – 1.0):
#   1. Count all quotation marks in the document.
#   2. Award 0.2 points if ANY curly quotes are present (evidence of work).
#   3. Award up to 0.8 additional points based on replacement coverage:
#        • Full 0.8 if **no** straight quotes remain.
#        • Otherwise proportional: 0.8 * (curly / total).
#   4. Cap and round the score to 1.0.
# ------------------------------------------------------------
# No points are given for file existence or loading (prerequisites only).
# ------------------------------------------------------------


def count_quotes_in_docx(file_path):
    """Return counts of straight and curly quotes in a DOCX file."""
    straight_double = straight_single = curly_double = curly_single = 0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Could not load document: {e}")
        return None

    # Scan paragraphs
    for p in doc.paragraphs:
        text = p.text
        straight_double += text.count('"')
        straight_single += text.count("'")
        curly_double += text.count('\u201C') + text.count('\u201D')
        curly_single += text.count('\u2018') + text.count('\u2019')

    # Scan tables (Writer often stores interview text in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                straight_double += text.count('"')
                straight_single += text.count("'")
                curly_double += text.count('\u201C') + text.count('\u201D')
                curly_single += text.count('\u2018') + text.count('\u2019')

    return {
        'straight_double': straight_double,
        'straight_single': straight_single,
        'curly_double': curly_double,
        'curly_single': curly_single,
    }


def compute_reward(file_path):
    """Compute progressive reward (0.0 – 1.0) for quote replacement."""
    if not os.path.exists(file_path):
        print(f"✗ File does not exist: {file_path}")
        return 0.0

    counts = count_quotes_in_docx(file_path)
    if counts is None:
        return 0.0  # Loading failed

    sd = counts['straight_double']
    ss = counts['straight_single']
    cd = counts['curly_double']
    cs = counts['curly_single']

    total_ascii = sd + ss
    total_curly = cd + cs
    total_quotes = total_ascii + total_curly

    print("Quote counts:")
    print(f"  Straight double (\u0022): {sd}")
    print(f"  Straight single (\u0027): {ss}")
    print(f"  Curly double (“ and ”): {cd}")
    print(f"  Curly single (‘ and ’): {cs}")
    print(f"  Total ASCII straight quotes: {total_ascii}")
    print(f"  Total curly quotes: {total_curly}")

    if total_quotes == 0:
        print("✗ No quotation marks found – cannot evaluate task")
        return 0.0

    score = 0.0

    # 1) Presence of any curly quotes (0.2 points)
    if total_curly > 0:
        score += 0.2
        print("✓ Curly quotes detected (0.2 points)")
    else:
        print("✗ No curly quotes detected (0 points)")

    # 2) Replacement coverage (up to 0.8 points)
    if total_ascii == 0 and total_curly > 0:
        score += 0.8
        print("✓ All straight quotes replaced (0.8 points)")
    else:
        coverage_ratio = total_curly / total_quotes  # 0–1
        coverage_points = 0.8 * coverage_ratio
        score += coverage_points
        print(f"Partial replacement: {coverage_ratio:.2%} curly → {coverage_points:.2f}/0.8 points")

    final_score = round(min(score, 1.0), 4)
    print(f"Total Score: {final_score}")
    return final_score


def verify_task():
    """Main entry point for the evaluation platform."""
    file_path = (
        '/home/user/'
        'i_just_dumped_a_12_page_interview_into_libreoffice_writer_and_every_single_quotation_mark_is_a_plain.docx'
    )
    reward = compute_reward(file_path)
    print(f"REWARD: {reward}")
    return reward


if __name__ == '__main__':
    verify_task()

