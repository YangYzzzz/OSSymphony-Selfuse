"""
Reward Script: Change page orientation to landscape
Task ID: writer_legal_009
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Orientation enum is LANDSCAPE
  Component 2 (0.3): Page width > page height (landscape dimensions)
  Component 3 (0.3): Content preserved (paragraphs and table intact)
"""

import os
from docx import Document
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_009'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section
    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: Orientation enum is LANDSCAPE (0.4 points)
    # This checks that the orientation property was explicitly set to landscape.
    # Initial: PORTRAIT -> FAIL; Golden: LANDSCAPE -> PASS
    try:
        orientation = section.orientation
        if orientation == WD_ORIENT.LANDSCAPE:
            print(f"PASS: Component 1 — Orientation is LANDSCAPE ({orientation}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected LANDSCAPE, found {orientation}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Page width > page height (landscape dimensions) (0.3 points)
    # In portrait: width=8.5in < height=11in -> FAIL
    # In landscape: width=11in > height=8.5in -> PASS
    try:
        pw = section.page_width
        ph = section.page_height
        if pw is not None and ph is not None and pw > ph:
            print(f"PASS: Component 2 — Width ({pw}) > Height ({ph}), landscape dimensions confirmed (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Width ({pw}) <= Height ({ph}), not landscape dimensions")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Content preserved — paragraphs and table intact (0.3 points)
    # The document must still have its content after orientation change.
    # Initial has 27 paragraphs and 1 table — golden should have at least
    # similar content. We check for the title and table existence.
    # This component is anchored to the orientation change: we only award
    # points if BOTH landscape orientation AND content are intact.
    try:
        # Check title paragraph exists
        title_found = any("Contract Comparison" in p.text for p in doc.paragraphs)
        # Check table exists with expected structure (at least 1 table with rows)
        table_found = len(doc.tables) >= 1 and len(doc.tables[0].rows) >= 2

        if section.orientation == WD_ORIENT.LANDSCAPE and title_found and table_found:
            print(f"PASS: Component 3 — Landscape orientation with content preserved: title found, table with {len(doc.tables[0].rows)} rows (0.3 pts)")
            total_score += 0.3
        else:
            if section.orientation != WD_ORIENT.LANDSCAPE:
                print(f"FAIL: Component 3 — Not in landscape mode, cannot award content preservation points")
            else:
                print(f"FAIL: Component 3 — Content missing: title_found={title_found}, table_found={table_found}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
