"""
Initial Setup: Contract comparison analysis document in portrait orientation
Task ID: writer_legal_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Ensure portrait orientation explicitly
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('Contract Comparison Analysis', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by: Whitfield & Associates Legal Consulting')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run = subtitle.add_run('\nDate: March 28, 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document provides a comprehensive side-by-side comparison of the '
        'proposed vendor contracts submitted by three shortlisted service providers '
        'for the enterprise cloud migration project. The analysis covers key '
        'contractual terms including liability caps, indemnification clauses, '
        'service level agreements (SLAs), termination provisions, and pricing '
        'structures. Our recommendation is based on a balanced assessment of risk '
        'exposure, cost-effectiveness, and operational flexibility.'
    )

    # Background
    doc.add_heading('2. Background', level=1)
    doc.add_paragraph(
        'Meridian Technologies Inc. issued a Request for Proposal (RFP-2026-0041) '
        'on January 15, 2026, seeking qualified vendors to manage the migration of '
        'on-premise infrastructure to a hybrid cloud environment. After an initial '
        'screening of twelve proposals, three vendors were shortlisted for detailed '
        'contract review:'
    )
    doc.add_paragraph('Apex Cloud Solutions — Proposal dated February 3, 2026', style='List Bullet')
    doc.add_paragraph('Nimbus Infrastructure Partners — Proposal dated February 7, 2026', style='List Bullet')
    doc.add_paragraph('SkyBridge Digital Services — Proposal dated February 10, 2026', style='List Bullet')

    # Comparison Table heading
    doc.add_heading('3. Contract Terms Comparison', level=1)
    doc.add_paragraph(
        'The following table summarizes the critical terms across all three vendor '
        'contracts. A detailed narrative analysis of each term follows in Section 4.'
    )

    # Comparison Table
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'

    headers = ['Contract Term', 'Apex Cloud Solutions', 'Nimbus Infrastructure', 'SkyBridge Digital']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['Contract Duration', '36 months', '24 months (renewable)', '36 months with 12-month exit clause'],
        ['Total Contract Value', '$2,450,000', '$1,890,000', '$2,180,000'],
        ['Liability Cap', '2x annual fees', '1x annual fees', '1.5x annual fees'],
        ['Uptime SLA', '99.95%', '99.9%', '99.99%'],
        ['SLA Penalty (per breach)', '5% monthly credit', '3% monthly credit', '10% monthly credit'],
        ['Termination Notice', '90 days', '60 days', '120 days'],
        ['Data Portability', 'Full export within 30 days', 'Export within 45 days, $15K fee', 'Full export within 15 days, no fee'],
        ['IP Ownership', 'Client retains all IP', 'Shared IP on custom modules', 'Client retains all IP'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val
            for run in table.cell(r, c).paragraphs[0].runs:
                run.font.size = Pt(9)

    # Detailed Analysis
    doc.add_heading('4. Detailed Analysis', level=1)

    doc.add_heading('4.1 Contract Duration and Flexibility', level=2)
    doc.add_paragraph(
        'Apex Cloud Solutions proposes a standard 36-month term with no early '
        'termination clause, which presents a moderate lock-in risk. Nimbus offers '
        'the shortest initial commitment at 24 months with automatic renewal, '
        'providing greater flexibility. SkyBridge\'s 36-month term is mitigated by '
        'a 12-month exit clause, allowing Meridian to reassess after the first year '
        'of service delivery.'
    )

    doc.add_heading('4.2 Pricing and Value Assessment', level=2)
    doc.add_paragraph(
        'While Nimbus presents the lowest total contract value at $1,890,000, their '
        'shorter term means the annualized cost ($945,000/year) exceeds that of '
        'Apex ($816,667/year) and SkyBridge ($726,667/year). When normalized to a '
        '36-month basis, SkyBridge offers the most competitive pricing, though their '
        'premium SLA commitments partially justify the higher absolute cost relative '
        'to Nimbus.'
    )

    doc.add_heading('4.3 Liability and Indemnification', level=2)
    doc.add_paragraph(
        'The liability cap is a critical risk factor. Apex\'s 2x annual fees cap '
        'provides the strongest protection for Meridian in the event of a major '
        'service failure. Nimbus\'s 1x cap is the most restrictive and may be '
        'insufficient for enterprise-grade risk tolerance. SkyBridge\'s 1.5x cap '
        'represents a reasonable middle ground. All three vendors include standard '
        'mutual indemnification for third-party IP claims.'
    )

    doc.add_heading('4.4 Service Level Agreements', level=2)
    doc.add_paragraph(
        'SkyBridge offers the highest uptime guarantee at 99.99%, backed by the '
        'most aggressive penalty structure (10% monthly credit per breach). This '
        'translates to approximately 52.6 minutes of allowable downtime per year. '
        'Apex\'s 99.95% SLA allows roughly 4.38 hours of annual downtime, while '
        'Nimbus\'s 99.9% permits up to 8.76 hours. For mission-critical workloads, '
        'the SkyBridge SLA provides the strongest guarantees.'
    )

    # Recommendation
    doc.add_heading('5. Recommendation', level=1)
    doc.add_paragraph(
        'Based on the comparative analysis, we recommend proceeding with contract '
        'negotiations with SkyBridge Digital Services. While not the lowest-cost '
        'option, SkyBridge offers the best combination of competitive pricing '
        '(when annualized), superior SLA commitments, full data portability at no '
        'additional cost, and complete IP ownership retention. The 12-month exit '
        'clause provides adequate flexibility to mitigate long-term commitment risk.'
    )

    doc.add_paragraph(
        'We further recommend negotiating the following modifications to the '
        'SkyBridge contract before execution:'
    )
    doc.add_paragraph('Reduce the termination notice period from 120 days to 90 days', style='List Number')
    doc.add_paragraph('Include a price escalation cap of 3% per annum for renewal terms', style='List Number')
    doc.add_paragraph('Add a dedicated account manager provision with quarterly business reviews', style='List Number')
    doc.add_paragraph('Incorporate a benchmarking clause allowing price adjustment based on market rates', style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
