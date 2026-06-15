"""
Reward Script: Volunteer sign-up sheet for Community Garden Project
Task ID: writer_wf_047
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Title "Community Garden Project" — 16pt, bold, centered
  Component 2 (0.25): Event details — dates, location, coordinator name, phone
  Component 3 (0.30): Sign-up table — 6 columns with correct headers, 16 rows (1 header + 15 blank)
  Component 4 (0.20): Bottom notes section — requirements/what to bring
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_047'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    tables = doc.tables

    # Component 1: Title "Community Garden Project" — 16pt, bold, centered (0.25 points)
    try:
        title_found = False
        for para in paragraphs:
            text = para.text.strip().lower()
            if 'community garden project' in text:
                # Check centered
                is_centered = (para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                # Check bold and 16pt in runs
                has_bold = False
                has_16pt = False
                for run in para.runs:
                    if run.font.bold:
                        has_bold = True
                    if run.font.size and abs(run.font.size.pt - 16.0) < 0.5:
                        has_16pt = True

                if is_centered and has_bold and has_16pt:
                    print(f"PASS: Component 1 — Title 'Community Garden Project' is 16pt, bold, centered (0.25 pts)")
                    total_score += 0.25
                    title_found = True
                else:
                    print(f"FAIL: Component 1 — Title found but: centered={is_centered}, bold={has_bold}, 16pt={has_16pt}")
                    title_found = True
                break
        if not title_found:
            print(f"FAIL: Component 1 — Title 'Community Garden Project' not found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Event details — dates, location, coordinator, phone (0.25 points)
    # Each sub-check worth 0.0625 (4 sub-checks)
    try:
        all_text = '\n'.join(p.text.lower() for p in paragraphs)
        detail_score = 0.0

        # Check dates info present
        if 'date' in all_text and any(month in all_text for month in ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']):
            detail_score += 0.0625
            print("  PASS: Dates info found")
        else:
            print("  FAIL: Dates info not found")

        # Check location present
        if 'location' in all_text:
            detail_score += 0.0625
            print("  PASS: Location info found")
        else:
            print("  FAIL: Location info not found")

        # Check coordinator name present
        if 'coordinator' in all_text:
            detail_score += 0.0625
            print("  PASS: Coordinator info found")
        else:
            print("  FAIL: Coordinator info not found")

        # Check phone present
        if 'phone' in all_text:
            detail_score += 0.0625
            print("  PASS: Phone info found")
        else:
            print("  FAIL: Phone info not found")

        if detail_score > 0:
            print(f"PASS: Component 2 — Event details ({detail_score} pts)")
            total_score += detail_score
        else:
            print(f"FAIL: Component 2 — No event details found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Sign-up table — 6 columns with correct headers, 16 rows (0.30 points)
    try:
        if len(tables) == 0:
            print(f"FAIL: Component 3 — No tables found in document")
        else:
            table = tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)

            table_score = 0.0

            # Sub-check 3a: Table has 6 columns (0.10 points)
            if num_cols == 6:
                print(f"  PASS: Table has 6 columns")
                table_score += 0.10
            else:
                print(f"  FAIL: Table has {num_cols} columns, expected 6")

            # Sub-check 3b: Header row has correct column names (0.10 points)
            expected_headers = ['name', 'email', 'phone', 'available dates', 'skills/equipment', 't-shirt size']
            actual_headers = [table.cell(0, c).text.strip().lower() for c in range(min(num_cols, 6))]
            header_matches = sum(1 for exp, act in zip(expected_headers, actual_headers) if exp in act or act in exp)
            if header_matches >= 5:
                print(f"  PASS: Header row has correct column names ({header_matches}/6 match)")
                table_score += 0.10
            else:
                print(f"  FAIL: Header row mismatch — expected {expected_headers}, found {actual_headers}")

            # Sub-check 3c: Table has 16 rows (1 header + 15 blank) (0.10 points)
            # Allow some tolerance: at least 14 rows total, up to 20
            if 14 <= num_rows <= 20:
                # Check that data rows (after header) are empty
                empty_rows = 0
                for ri in range(1, num_rows):
                    row_text = ''.join(table.cell(ri, c).text.strip() for c in range(min(num_cols, 6)))
                    if row_text == '':
                        empty_rows += 1
                if empty_rows >= 13:
                    print(f"  PASS: Table has {num_rows} rows with {empty_rows} empty sign-up rows")
                    table_score += 0.10
                else:
                    print(f"  FAIL: Only {empty_rows} empty rows out of {num_rows - 1} data rows")
            else:
                print(f"  FAIL: Table has {num_rows} rows, expected 14-20 (ideally 16)")

            total_score += table_score
            if table_score > 0:
                print(f"PASS: Component 3 — Sign-up table ({table_score} pts)")
            else:
                print(f"FAIL: Component 3 — Table structure incorrect")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Bottom notes section — requirements and what to bring (0.20 points)
    try:
        # Notes should appear AFTER the table in the document
        # Look for keywords related to requirements/what to bring
        notes_text = ''
        # Gather text from paragraphs that appear after at least some content
        found_table_context = False
        for para in paragraphs:
            text = para.text.strip().lower()
            if 'note' in text or 'important' in text or 'requirement' in text:
                found_table_context = True
            if found_table_context:
                notes_text += ' ' + text

        notes_score = 0.0

        # Check for notes header/section
        if 'note' in notes_text or 'important' in notes_text or 'requirement' in notes_text:
            notes_score += 0.10
            print("  PASS: Notes section header found")
        else:
            print("  FAIL: Notes section header not found")

        # Check for actual content about requirements/what to bring
        has_requirement_content = any(kw in notes_text for kw in ['bring', 'wear', 'shoe', 'glove', 'tool', 'water', 'age', 'weather', 'cancel'])
        if has_requirement_content:
            notes_score += 0.10
            print("  PASS: Notes contain requirement/what-to-bring content")
        else:
            print("  FAIL: Notes lack requirement/what-to-bring content")

        total_score += notes_score
        if notes_score > 0:
            print(f"PASS: Component 4 — Bottom notes ({notes_score} pts)")
        else:
            print(f"FAIL: Component 4 — No bottom notes found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
