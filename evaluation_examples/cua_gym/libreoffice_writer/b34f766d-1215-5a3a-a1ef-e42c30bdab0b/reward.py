"""
FINAL REWARD SCRIPT - SUCCESS
Task: LibreOffice Writer keeps chaining the numbers from my first list into the next one. My first list reads 1) Scope, 2) Timeline, 3) Budget. When I start a new numbered list two pages later for the section called “Risk Categories,” it jumps straight to 4) instead of 1). How do I tell Writer to restart that second list so it begins at 1) again—i.e., 1) Operational, 2) Financial, 3) Compliance?
Generated: 2025-09-10 18:49:33
Status: success
Model: azure-o3
Total Steps: 11
"""

import os
import re
from docx import Document


def extract_numbered_blocks(doc):
    """Extract contiguous blocks of manually numbered lines like '1) Text'.
    Returns a list of blocks, each block being a list of integers (the numbers)."""
    blocks = []
    current_block = []
    num_pattern = re.compile(r"^\s*(\d+)\)")  # matches leading number followed by a right parenthesis

    for para in doc.paragraphs:
        text = para.text.strip()
        match = num_pattern.match(text)
        if match:
            current_block.append(int(match.group(1)))
        else:
            if current_block:
                blocks.append(current_block)
                current_block = []

    # append last collected block if any
    if current_block:
        blocks.append(current_block)

    return blocks


def is_sequential_from_one(block):
    """Check if a block starts at 1 and increments sequentially by 1."""
    if not block or block[0] != 1:
        return False
    return all(num == idx for idx, num in enumerate(block, start=1))


def verify_writer_task(file_path):
    """Reward-script verifier for the LibreOffice Writer list-restart task.
    Scoring (progressive):
        0.5 points – first numbered list starts at 1 and is sequential
        0.5 points – a subsequent numbered list restarts at 1 and is sequential
        1.0 points – both conditions satisfied
    Returns a float between 0.0 and 1.0 and prints detailed diagnostics."""

    print(f"Verifying numbered-list restart in file: {file_path}")

    # ---------- prerequisite checks (no points awarded) ----------
    if not os.path.exists(file_path):
        print("✗ File not found – cannot verify task.")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
        print(f"✓ Loaded DOCX successfully (paragraphs: {len(doc.paragraphs)}) – no points awarded")
    except Exception as e:
        print(f"✗ Error loading DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- actual verification begins ----------
    blocks = extract_numbered_blocks(doc)
    print(f"Detected {len(blocks)} numbered block(s): {blocks}")

    if not blocks:
        print("✗ No manually-numbered lists detected. Task incomplete.")
        print("REWARD: 0.0")
        return 0.0

    score = 0.0

    # Requirement 1: first list correct
    if is_sequential_from_one(blocks[0]):
        print("✓ First numbered list starts at 1 and is sequential (+0.5)")
        score += 0.5
    else:
        print("✗ First numbered list is not correct (+0.0)")

    # Requirement 2: a later list restarts at 1 and is sequential
    if len(blocks) >= 2:
        if is_sequential_from_one(blocks[1]):
            print("✓ Second numbered list restarts at 1 and is sequential (+0.5)")
            score += 0.5
        else:
            print("✗ Second numbered list does not restart at 1 or is not sequential (+0.0)")
    else:
        print("✗ Fewer than two numbered lists found (+0.0)")

    final_score = min(score, 1.0)
    print(f"Total score: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# ----------------- script execution -----------------
if __name__ == "__main__":
    TARGET_PATH = "/home/user/libreoffice_writer_keeps_chaining_the_numbers_from_my_first_list_into_the_next_one_my_first_list_rea.docx"
    verify_writer_task(TARGET_PATH)

