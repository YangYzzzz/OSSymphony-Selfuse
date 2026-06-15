"""
Reward Script: Remove numbering from 'Note: This step is optional' paragraph
Task ID: writer_lec_024
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.5): Note paragraph style is not a list style (changed from List Number)
  - Component 2 (0.3): Note is de-listed AND subsequent list items still have List Number style
  - Component 3 (0.2): Note is de-listed AND note text is preserved intact
All components are gated on the note paragraph having been de-listed, ensuring
initial_env scores 0.0.
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_024'

NOTE_TEXT = 'Note: This step is optional'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def is_list_number_style(para):
    """Check if a paragraph has List Number style or numPr XML."""
    style_name = para.style.name if para.style else ''
    is_list = 'list' in style_name.lower() and 'number' in style_name.lower()

    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    pPr = para._element.find(f'{{{ns}}}pPr')
    has_numPr = (
        pPr is not None
        and pPr.find(f'{{{ns}}}numPr') is not None
    )

    return is_list or has_numPr


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the note paragraph by its text content
    note_para = None
    note_idx = None
    for i, para in enumerate(doc.paragraphs):
        if NOTE_TEXT in para.text:
            note_para = para
            note_idx = i
            break

    if note_para is None:
        print("CRITICAL: Could not find paragraph containing 'Note: This step is optional'")
        print("REWARD: 0.0")
        return 0.0

    style_name = note_para.style.name if note_para.style else ''
    note_is_delisted = not is_list_number_style(note_para)

    print(f"INFO: Found note paragraph at index {note_idx}: text={note_para.text!r}, style={style_name!r}, delisted={note_is_delisted}")

    # Component 1: Note paragraph is NOT a numbered list item (0.5 points)
    # This is the primary change: style should no longer be 'List Number'.
    # FAILS on initial (style='List Number'), PASSES on golden (style='Normal').
    try:
        if note_is_delisted:
            print(f"PASS: Component 1 - Note paragraph is not numbered (style={style_name!r}) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 - Note paragraph still appears numbered (style={style_name!r})")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Note is de-listed AND subsequent list items retain 'List Number' style (0.3 points)
    # Gated on note_is_delisted so it scores 0 on initial_env.
    # Verifies the agent didn't break the remaining numbered list.
    try:
        if note_is_delisted:
            post_note_list_items = []
            for j in range(note_idx + 1, len(doc.paragraphs)):
                para = doc.paragraphs[j]
                if not is_list_number_style(para):
                    break
                post_note_list_items.append(para)

            if len(post_note_list_items) >= 3:
                print(f"PASS: Component 2 - Note de-listed AND {len(post_note_list_items)} subsequent items retain list style (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 - Expected >=3 list items after note, found {len(post_note_list_items)}")
        else:
            print(f"FAIL: Component 2 - Note not de-listed, skipping subsequent list check")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Note is de-listed AND note text is preserved (0.2 points)
    # Gated on note_is_delisted so it scores 0 on initial_env.
    # Verifies the agent didn't delete/corrupt the note text.
    try:
        if note_is_delisted:
            actual_text = note_para.text.strip()
            if NOTE_TEXT in actual_text:
                print(f"PASS: Component 3 - Note de-listed AND text preserved: {actual_text!r} (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 - Note text changed: {actual_text!r}")
        else:
            print(f"FAIL: Component 3 - Note not de-listed, skipping text preservation check")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved state before verifying
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
