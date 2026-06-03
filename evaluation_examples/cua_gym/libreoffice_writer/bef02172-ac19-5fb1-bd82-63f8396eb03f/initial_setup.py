"""
Initial Setup: Remove numbering from a note paragraph incorrectly in a numbered list
Task ID: writer_lec_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Setting Up Your Home Network Router', level=1)

    # Introduction paragraph
    intro = doc.add_paragraph(
        'Follow these steps carefully to configure your new wireless router. '
        'Make sure you have your ISP credentials ready before starting.'
    )

    # Numbered list of steps (8 items total: 7 real steps + 1 note incorrectly in the list)
    # Items 1-4
    steps_before_note = [
        'Unbox the router and connect the power adapter to the unit. '
        'Wait for the power LED to turn solid green before proceeding.',
        'Connect an Ethernet cable from the modem\'s LAN port to the router\'s WAN port. '
        'The WAN port is usually colored differently (blue or yellow).',
        'On your computer, open a web browser and navigate to 192.168.1.1 to access '
        'the router\'s administration panel.',
        'Log in using the default credentials printed on the label underneath the router. '
        'The default username is usually "admin" and the password is on the sticker.',
    ]

    for step_text in steps_before_note:
        doc.add_paragraph(step_text, style='List Number')

    # The note paragraph - incorrectly included in the numbered list as item 5
    doc.add_paragraph(
        'Note: This step is optional',
        style='List Number'
    )

    # Items that come after the note (3 more steps, currently numbered 6, 7, 8)
    steps_after_note = [
        'Navigate to the Wireless Settings section and set your preferred network name (SSID). '
        'Choose a unique name that doesn\'t reveal personal information.',
        'Under Security Settings, select WPA3-Personal encryption and create a strong '
        'password with at least 12 characters including numbers and symbols.',
        'Click "Apply" or "Save Settings" to finalize your configuration. The router will '
        'reboot automatically, which takes approximately 60 seconds.',
    ]

    for step_text in steps_after_note:
        doc.add_paragraph(step_text, style='List Number')

    # Closing paragraph
    closing = doc.add_paragraph(
        'Once the router has rebooted, try connecting a device to the new wireless network '
        'using the SSID and password you configured. If you experience any issues, '
        'consult the troubleshooting section in the user manual.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
