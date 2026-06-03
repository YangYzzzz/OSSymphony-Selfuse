"""
Initial Setup: Multi-level numbered list with items at mixed levels
Task ID: writer_list_009
Domain: libreoffice_writer

Creates report_outline.docx with 7 items:
  - Items 1, 2, 4, 6, 7 at level 1 (List Number)
  - Items 3 ("Methodology") and 5 ("Analysis and Findings") at level 2 (List Number 2)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_list_009'
OUTPUT = f'{WORKDIR}/Desktop/report_outline.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_list_level(paragraph, level):
    """
    Set the numbering level for a list paragraph using XML manipulation.
    level: 0-based (0 = level 1, 1 = level 2)
    """
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pPr.insert(0, numPr)

    ilvl = numPr.find(qn('w:ilvl'))
    if ilvl is None:
        ilvl = OxmlElement('w:ilvl')
        numPr.insert(0, ilvl)
    ilvl.set(qn('w:val'), str(level))


def create_initial():
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document()

    # Define the items and their levels (0-based: 0=level1, 1=level2)
    items = [
        ("Introduction", 0),
        ("Background Research", 0),
        ("Methodology", 1),          # level 2 sub-item
        ("Data Collection", 0),
        ("Analysis and Findings", 1), # level 2 sub-item
        ("Recommendations", 0),
        ("Conclusion", 0),
    ]

    for text, level in items:
        if level == 0:
            para = doc.add_paragraph(text, style='List Number')
        else:
            para = doc.add_paragraph(text, style='List Number 2')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
