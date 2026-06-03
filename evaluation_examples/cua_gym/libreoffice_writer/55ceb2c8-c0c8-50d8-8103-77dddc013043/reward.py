"""
Reward Script: Update TOC to include new Chapter 6 entries
Task ID: writer_rm_065
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): TOC contains "Chapter 6: New Features" entry
  Component 2 (0.3): TOC contains all three sub-heading entries (6.1, 6.2, 6.3)
  Component 3 (0.2): TOC entry count increased (>= 24 entries vs initial 20)
  Component 4 (0.2): Chapter 6 body content is intact with correct headings
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_065'


def get_toc_entries(doc):
    """
    Extract TOC entries: paragraphs between 'Table of Contents' heading
    and the first non-TOC content (next Heading or empty paragraph sequence).
    TOC entries are Normal-style paragraphs with tab-separated page numbers.
    """
    toc_entries = []
    in_toc = False
    blank_count = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''

        if style_name == 'Heading 1' and 'Table of Contents' in para.text:
            in_toc = True
            continue

        if in_toc:
            text = para.text.strip()
            # End of TOC: hit a Heading style (body content starts)
            if style_name.startswith('Heading'):
                break
            # Track consecutive blanks; 2+ blanks signals end of TOC
            if not text:
                blank_count += 1
                if blank_count >= 2:
                    break
                continue
            else:
                blank_count = 0

            # TOC entry: has text with a tab (page number separator)
            if '\t' in text:
                toc_entries.append(text)

    return toc_entries


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract TOC entries
    toc_entries = get_toc_entries(doc)
    toc_text_lower = [e.lower() for e in toc_entries]
    print(f"INFO: Found {len(toc_entries)} TOC entries")
    for i, entry in enumerate(toc_entries):
        print(f"  TOC[{i}]: {entry}")

    # Component 1: TOC contains "Chapter 6: New Features" entry (0.3 points)
    # This FAILS on initial (no Chapter 6 in TOC) -> PASSES on golden
    try:
        ch6_found = any('chapter 6' in e and 'new features' in e for e in toc_text_lower)
        if ch6_found:
            print(f"PASS: Component 1 - 'Chapter 6: New Features' found in TOC (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 - 'Chapter 6: New Features' NOT found in TOC")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: TOC contains all three Heading 2 sub-entries (0.3 points)
    # 6.1 AI-Assisted Code Review, 6.2 Real-Time Collaboration Tools, 6.3 Automated Dependency Management
    # This FAILS on initial -> PASSES on golden
    try:
        sub_entries = [
            ('6.1', 'ai-assisted code review'),
            ('6.2', 'real-time collaboration tools'),
            ('6.3', 'automated dependency management'),
        ]
        found_subs = 0
        for num, title_fragment in sub_entries:
            entry_found = any(num in e and title_fragment in e for e in toc_text_lower)
            if entry_found:
                found_subs += 1
                print(f"  SUB-ENTRY: {num} found in TOC")
            else:
                print(f"  SUB-ENTRY: {num} ({title_fragment}) NOT found in TOC")

        if found_subs == 3:
            print(f"PASS: Component 2 - All 3 sub-headings found in TOC (0.3 pts)")
            total_score += 0.3
        elif found_subs > 0:
            partial = round(0.3 * found_subs / 3, 2)
            print(f"PARTIAL: Component 2 - {found_subs}/3 sub-headings found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - No Chapter 6 sub-headings found in TOC")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: TOC entry count increased to >= 24 (initial has 20, golden has 24) (0.2 points)
    # This FAILS on initial (20 entries) -> PASSES on golden (24 entries)
    try:
        if len(toc_entries) >= 24:
            print(f"PASS: Component 3 - TOC has {len(toc_entries)} entries (>= 24) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 - TOC has {len(toc_entries)} entries, expected >= 24")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Chapter 6 body content intact (headings present in body) (0.2 points)
    # This is a compound check: Chapter 6 heading exists in body AND TOC was updated (anchored to change)
    # We check that the body still has Chapter 6 content AND the TOC references it
    try:
        body_headings = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading') and 'new features' in para.text.lower():
                body_headings.append(para.text)
            elif style_name.startswith('Heading') and para.text.strip().startswith('6.'):
                body_headings.append(para.text)

        ch6_body_ok = len(body_headings) >= 4  # Chapter 6 heading + 3 sub-headings
        ch6_in_toc = any('chapter 6' in e and 'new features' in e for e in toc_text_lower)

        # Only award points if BOTH body is intact AND TOC was updated
        if ch6_body_ok and ch6_in_toc:
            print(f"PASS: Component 4 - Chapter 6 body intact ({len(body_headings)} headings) AND TOC updated (0.2 pts)")
            total_score += 0.2
        elif ch6_body_ok and not ch6_in_toc:
            print(f"FAIL: Component 4 - Body has Chapter 6 but TOC not updated")
        else:
            print(f"FAIL: Component 4 - Chapter 6 body headings missing (found {len(body_headings)})")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
