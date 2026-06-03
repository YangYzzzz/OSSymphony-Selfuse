"""
Initial Setup: Master document with outdated TOC missing new Chapter 6
Task ID: writer_rm_065
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_065'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_field(doc):
    """
    Add a TOC field to the document. This creates the TOC field codes
    that LibreOffice will recognize and can update.
    """
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles['Normal']
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add TOC field - begin
    run_begin = paragraph.add_run()
    fld_char_begin = run_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_begin._element.append(fld_char_begin)

    # TOC instruction
    run_instr = paragraph.add_run()
    instr_text = run_instr._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr_text.text = r' TOC \o "1-3" \h \z \u '
    run_instr._element.append(instr_text)

    # Separate
    run_sep = paragraph.add_run()
    fld_char_sep = run_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run_sep._element.append(fld_char_sep)

    return paragraph


def add_toc_entry(doc, text, level=1, page_num="1"):
    """Add a static TOC entry line (cached display text for the outdated TOC)."""
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    if level == 1:
        para.paragraph_format.left_indent = Inches(0)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
    else:
        para.paragraph_format.left_indent = Inches(0.3 * (level - 1))
        run = para.add_run(text)
        run.font.size = Pt(10)
    # Add tab and page number
    tab_run = para.add_run(f'\t{page_num}')
    tab_run.font.size = Pt(10) if level > 1 else Pt(11)
    return para


def close_toc_field(doc):
    """Close the TOC field."""
    paragraph = doc.add_paragraph()
    run_end = paragraph.add_run()
    fld_char_end = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_end._element.append(fld_char_end)
    return paragraph


def create_initial():
    doc = Document()

    # ---- Document Title ----
    title = doc.add_heading('Technical Handbook - Software Development Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Version 3.2 — March 2025')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacing

    # ---- Table of Contents (OUTDATED - missing Chapter 6) ----
    toc_heading = doc.add_heading('Table of Contents', level=1)

    # Add TOC field
    add_toc_field(doc)

    # Static TOC entries (only Chapters 1-5, deliberately missing Chapter 6)
    toc_entries = [
        ("Chapter 1: Introduction to Development", 1, "3"),
        ("1.1 Purpose and Scope", 2, "3"),
        ("1.2 Target Audience", 2, "4"),
        ("1.3 Document Conventions", 2, "5"),
        ("Chapter 2: Environment Setup", 1, "7"),
        ("2.1 Hardware Requirements", 2, "7"),
        ("2.2 Software Prerequisites", 2, "8"),
        ("2.3 Configuration Steps", 2, "10"),
        ("Chapter 3: Coding Standards", 1, "13"),
        ("3.1 Naming Conventions", 2, "13"),
        ("3.2 Code Structure", 2, "15"),
        ("3.3 Documentation Requirements", 2, "17"),
        ("Chapter 4: Testing Framework", 1, "20"),
        ("4.1 Unit Testing", 2, "20"),
        ("4.2 Integration Testing", 2, "23"),
        ("4.3 Performance Testing", 2, "25"),
        ("Chapter 5: Deployment Pipeline", 1, "28"),
        ("5.1 CI/CD Configuration", 2, "28"),
        ("5.2 Staging Environment", 2, "30"),
        ("5.3 Production Release", 2, "32"),
    ]

    for text, level, page in toc_entries:
        add_toc_entry(doc, text, level, page)

    close_toc_field(doc)

    doc.add_page_break()

    # ---- Chapter 1: Introduction to Development ----
    doc.add_heading('Chapter 1: Introduction to Development', level=1)
    doc.add_heading('1.1 Purpose and Scope', level=2)
    doc.add_paragraph(
        'This handbook serves as the definitive guide for software development practices '
        'within the organization. It covers essential topics from environment setup through '
        'deployment, ensuring consistency across all development teams.'
    )
    doc.add_heading('1.2 Target Audience', level=2)
    doc.add_paragraph(
        'This document is intended for software developers, DevOps engineers, QA specialists, '
        'and technical leads. Familiarity with basic programming concepts is assumed.'
    )
    doc.add_heading('1.3 Document Conventions', level=2)
    doc.add_paragraph(
        'Code examples are presented in monospace font. Important warnings are highlighted '
        'in bold text. Cross-references use hyperlinked section numbers.'
    )

    doc.add_page_break()

    # ---- Chapter 2: Environment Setup ----
    doc.add_heading('Chapter 2: Environment Setup', level=1)
    doc.add_heading('2.1 Hardware Requirements', level=2)
    doc.add_paragraph(
        'Development workstations require a minimum of 16GB RAM, a quad-core processor, '
        'and 256GB SSD storage. For machine learning projects, a dedicated GPU with at '
        'least 8GB VRAM is recommended.'
    )
    doc.add_heading('2.2 Software Prerequisites', level=2)
    doc.add_paragraph(
        'All developers must install the following: Python 3.10+, Node.js 18 LTS, '
        'Docker Desktop, Git 2.40+, and the approved IDE (VSCode or IntelliJ IDEA). '
        'Platform-specific installers are available on the internal portal.'
    )
    doc.add_heading('2.3 Configuration Steps', level=2)
    doc.add_paragraph(
        'After installing the prerequisites, run the bootstrap script located at '
        '/tools/setup/bootstrap.sh. This configures environment variables, SSH keys, '
        'and connects to the internal package registry.'
    )

    doc.add_page_break()

    # ---- Chapter 3: Coding Standards ----
    doc.add_heading('Chapter 3: Coding Standards', level=1)
    doc.add_heading('3.1 Naming Conventions', level=2)
    doc.add_paragraph(
        'Use camelCase for variables and functions, PascalCase for classes and interfaces, '
        'and SCREAMING_SNAKE_CASE for constants. Database table names use snake_case. '
        'Avoid abbreviations except for widely recognized ones (e.g., URL, API, ID).'
    )
    doc.add_heading('3.2 Code Structure', level=2)
    doc.add_paragraph(
        'Each module follows the standard directory layout: src/ for source files, '
        'tests/ for test suites, docs/ for documentation, and config/ for environment-specific '
        'settings. Maximum file length is 500 lines; exceeding this indicates a need for refactoring.'
    )
    doc.add_heading('3.3 Documentation Requirements', level=2)
    doc.add_paragraph(
        'All public APIs must have docstrings following the Google style guide. '
        'README files are required at each module root. Architecture Decision Records (ADRs) '
        'must be filed for significant design choices.'
    )

    doc.add_page_break()

    # ---- Chapter 4: Testing Framework ----
    doc.add_heading('Chapter 4: Testing Framework', level=1)
    doc.add_heading('4.1 Unit Testing', level=2)
    doc.add_paragraph(
        'Unit tests are written using pytest (Python) or Jest (JavaScript). '
        'Minimum code coverage is 80% for all modules. Tests must be deterministic '
        'and independent — no shared mutable state between test cases.'
    )
    doc.add_heading('4.2 Integration Testing', level=2)
    doc.add_paragraph(
        'Integration tests verify interactions between services using Docker Compose '
        'environments. Test databases are provisioned automatically and seeded with '
        'fixture data. Each integration test suite runs in under 5 minutes.'
    )
    doc.add_heading('4.3 Performance Testing', level=2)
    doc.add_paragraph(
        'Load testing uses Locust for API endpoints with a baseline of 1000 concurrent users. '
        'Response time P95 must be under 200ms for critical paths. Performance regression '
        'tests run nightly and block releases if thresholds are exceeded.'
    )

    doc.add_page_break()

    # ---- Chapter 5: Deployment Pipeline ----
    doc.add_heading('Chapter 5: Deployment Pipeline', level=1)
    doc.add_heading('5.1 CI/CD Configuration', level=2)
    doc.add_paragraph(
        'The CI/CD pipeline uses GitHub Actions with the following stages: lint, test, build, '
        'security scan, and deploy. All stages must pass before merging to the main branch. '
        'Pipeline definitions are stored in .github/workflows/.'
    )
    doc.add_heading('5.2 Staging Environment', level=2)
    doc.add_paragraph(
        'Staging mirrors production with reduced resource allocation. All feature branches '
        'are automatically deployed to staging upon PR creation. Staging deployments include '
        'synthetic monitoring and automated smoke tests.'
    )
    doc.add_heading('5.3 Production Release', level=2)
    doc.add_paragraph(
        'Production releases follow a blue-green deployment strategy. Canary releases '
        'handle 5% of traffic initially, scaling to 100% over 30 minutes if error rates '
        'remain below 0.1%. Rollback is automated if health checks fail.'
    )

    doc.add_page_break()

    # ---- Chapter 6: New Features (content exists but NOT in TOC) ----
    doc.add_heading('Chapter 6: New Features', level=1)
    doc.add_heading('6.1 AI-Assisted Code Review', level=2)
    doc.add_paragraph(
        'The new AI-assisted code review system integrates with pull requests to provide '
        'automated suggestions for code quality improvements, security vulnerabilities, '
        'and performance optimizations. Reviewers can accept, modify, or dismiss suggestions '
        'directly from the PR interface.'
    )
    doc.add_heading('6.2 Real-Time Collaboration Tools', level=2)
    doc.add_paragraph(
        'Live pair programming sessions are now supported through the integrated collaboration '
        'platform. Developers can share their IDE workspace with teammates, co-edit files '
        'simultaneously, and communicate via built-in voice channels without leaving the editor.'
    )
    doc.add_heading('6.3 Automated Dependency Management', level=2)
    doc.add_paragraph(
        'The dependency management bot automatically monitors all project dependencies for '
        'security patches and version updates. It creates PRs for critical updates within '
        '24 hours of CVE publication and schedules non-critical updates for the weekly '
        'maintenance window.'
    )

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
