"""
Initial Setup: APA-style References section with plain paragraphs (no hanging indent)
Task ID: writer_acad_062
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_062'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

REFERENCES = [
    "Anderson, T. R., & Kim, S. J. (2021). Machine learning approaches to predicting urban heat island effects in metropolitan areas. Journal of Environmental Science & Technology, 45(3), 234-251. https://doi.org/10.1016/j.jest.2021.03.015",
    "Baker, L. M., Owens, D. R., & Patel, N. (2019). Longitudinal effects of mindfulness-based stress reduction on cortisol levels in healthcare workers. Psychoneuroendocrinology, 108, 42-50. https://doi.org/10.1016/j.psyneuen.2019.06.003",
    "Chen, W., Alvarez, R. F., & Nakamura, K. (2022). Structural integrity of additively manufactured titanium alloys under cyclic loading conditions. Materials Science and Engineering: A, 831, 142207. https://doi.org/10.1016/j.msea.2021.142207",
    "Davis, E. P., & Hernandez, M. L. (2020). The role of parental scaffolding in early childhood executive function development: A meta-analysis. Developmental Psychology, 56(8), 1503-1518. https://doi.org/10.1037/dev0000987",
    "Fischer, G., Lemoine, B., & Tran, Q. H. (2023). Carbon capture efficiency in deep saline aquifers: Comparative modeling of three injection strategies. Energy & Environmental Science, 16(2), 678-695. https://doi.org/10.1039/D2EE03401K",
    "Gonzalez, A. R., Whitfield, S. T., & Bauer, J. (2018). Community-based participatory research in Indigenous health: Ethical frameworks and practical challenges. American Journal of Public Health, 108(S2), S123-S130. https://doi.org/10.2105/AJPH.2017.304036",
    "Harrison, P. J., Liu, X., & Okonkwo, C. E. (2021). Quantum entanglement verification in noisy intermediate-scale quantum processors. Physical Review Letters, 127(15), 150502. https://doi.org/10.1103/PhysRevLett.127.150502",
    "Inoue, Y., & Castellano, D. (2022). Trade liberalization and income inequality in developing economies: Evidence from panel data analysis of 47 countries. World Development, 158, 105978. https://doi.org/10.1016/j.worlddev.2022.105978",
    "Johansson, K., Eriksson, M., & Sundstrom, A. (2020). Biomechanical analysis of anterior cruciate ligament reconstruction using hamstring tendon versus patellar tendon autografts. Journal of Orthopaedic Research, 38(5), 1102-1112. https://doi.org/10.1002/jor.24567",
    "Kumar, V., Robinson, E. A., & de Souza, F. (2019). Deep reinforcement learning for autonomous vehicle navigation in unstructured environments. IEEE Transactions on Intelligent Transportation Systems, 20(11), 4062-4075. https://doi.org/10.1109/TITS.2019.2897654",
    "Lee, S. H., Martinez, C., & O'Brien, W. F. (2023). CRISPR-Cas9 gene editing for sickle cell disease: Long-term follow-up data from phase II clinical trials. The New England Journal of Medicine, 388(12), 1089-1101. https://doi.org/10.1056/NEJMoa2215734",
    "Moretti, A., Singh, R., & Zhao, L. (2021). Socioeconomic determinants of food insecurity during the COVID-19 pandemic: A cross-sectional study of 12 European nations. The Lancet Regional Health - Europe, 4, 100085. https://doi.org/10.1016/j.lanepe.2021.100085",
    "Nguyen, T. D., Park, J. W., & Hoffmann, R. (2020). Electrochemical reduction of CO2 to formate on bismuth-based catalysts in aqueous electrolytes. ACS Catalysis, 10(18), 10726-10740. https://doi.org/10.1021/acscatal.0c02615",
    "Olsen, B. K., Yamamoto, H., & Fraser, C. (2022). The impact of remote work on organizational culture and employee well-being: A mixed-methods longitudinal study. Journal of Applied Psychology, 107(9), 1534-1552. https://doi.org/10.1037/apl0001045",
    "Petrov, I. A., Williams, G. N., & Abadi, S. (2018). Glacier mass balance trends in the Karakoram Range: Reconciling in situ measurements with satellite gravimetry data. The Cryosphere, 12(10), 3223-3240. https://doi.org/10.5194/tc-12-3223-2018",
]


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("The Effects of Urban Green Spaces on Mental Health Outcomes: A Systematic Review", level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph(
        "This systematic review examines the relationship between urban green spaces and mental health "
        "outcomes across 47 peer-reviewed studies published between 2015 and 2023. Results indicate that "
        "proximity to and engagement with green spaces is associated with reduced symptoms of depression "
        "(d = 0.42, 95% CI [0.31, 0.53]), anxiety (d = 0.38, 95% CI [0.27, 0.49]), and perceived stress "
        "(d = 0.35, 95% CI [0.22, 0.48]). Moderator analyses revealed that the type of green space, "
        "duration of exposure, and socioeconomic context significantly influenced effect sizes. "
        "Implications for urban planning policy and future research directions are discussed."
    )

    # Introduction section (abbreviated)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "Urbanization continues to accelerate globally, with projections indicating that 68% of the world's "
        "population will reside in urban areas by 2050 (United Nations, 2018). This demographic shift has "
        "coincided with rising rates of mental health disorders, prompting researchers to investigate "
        "environmental factors that may buffer against psychological distress (Anderson & Kim, 2021). "
        "Among these factors, urban green spaces—defined as publicly accessible areas with natural vegetation "
        "including parks, gardens, greenways, and urban forests—have received increasing scholarly attention."
    )
    doc.add_paragraph(
        "Theoretical frameworks linking green spaces to mental health draw from attention restoration theory "
        "(Kaplan, 1995), stress reduction theory (Ulrich et al., 1991), and more recent biophilic design "
        "principles (Fischer et al., 2023). These frameworks converge on the idea that natural environments "
        "facilitate cognitive recovery, reduce physiological stress markers, and promote social cohesion "
        "(Davis & Hernandez, 2020; Gonzalez et al., 2018)."
    )

    # Methods section (abbreviated)
    doc.add_heading("Methods", level=2)
    doc.add_paragraph(
        "A systematic search was conducted across PubMed, PsycINFO, Web of Science, and Scopus databases "
        "using a combination of MeSH terms and free-text keywords. Studies were eligible for inclusion if "
        "they (a) employed a quantitative research design, (b) measured at least one standardized mental "
        "health outcome, (c) assessed exposure to urban green spaces, and (d) were published in English "
        "in a peer-reviewed journal between January 2015 and December 2023."
    )

    # References section - plain paragraphs, NO hanging indent
    doc.add_heading("References", level=2)

    for ref in REFERENCES:
        para = doc.add_paragraph(ref)
        # Explicitly ensure no indentation
        para.paragraph_format.left_indent = None
        para.paragraph_format.first_line_indent = None
        # Set font to Times New Roman 12pt for academic style
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
