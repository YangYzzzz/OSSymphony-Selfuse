"""
Initial Setup: Create a formal letter document with an address block
Task ID: writer_lec_056
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard US Letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Sender Information (top right) ---
    sender_para = doc.add_paragraph()
    sender_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sender_para.paragraph_format.space_after = Pt(0)
    run = sender_para.add_run("Horizon Creative Solutions")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    sender_addr1 = doc.add_paragraph()
    sender_addr1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sender_addr1.paragraph_format.space_after = Pt(0)
    run = sender_addr1.add_run("1200 Innovation Drive, Suite 300")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    sender_addr2 = doc.add_paragraph()
    sender_addr2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sender_addr2.paragraph_format.space_after = Pt(0)
    run = sender_addr2.add_run("San Francisco, CA 94107")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    sender_phone = doc.add_paragraph()
    sender_phone.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sender_phone.paragraph_format.space_after = Pt(0)
    run = sender_phone.add_run("Tel: (415) 555-0198")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    sender_email = doc.add_paragraph()
    sender_email.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sender_email.paragraph_format.space_after = Pt(12)
    run = sender_email.add_run("info@horizoncreative.com")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.add_run("March 28, 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # --- Recipient Address Block ---
    addr_lines = [
        "Ms. Linda Park",
        "Marketing Director",
        "BrightStar Media",
        "456 Sunset Blvd",
        "Los Angeles, CA 90028",
    ]
    for i, line in enumerate(addr_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if i == len(addr_lines) - 1:
            p.paragraph_format.space_after = Pt(12)
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    # --- Salutation ---
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(12)
    run = salutation.add_run("Dear Ms. Park,")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # --- Body paragraphs ---
    body1_text = (
        "Thank you for taking the time to meet with our team last Thursday to discuss "
        "the upcoming digital marketing campaign for the Q3 product launch. We were "
        "impressed by BrightStar Media's creative portfolio and believe your agency is "
        "an excellent fit for our brand vision."
    )
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(12)
    run = body1.add_run(body1_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    body2_text = (
        "As discussed, we would like to move forward with the integrated social media "
        "strategy you proposed, including the influencer outreach program and the "
        "targeted video content series. Our budget allocation for this initiative is "
        "$125,000, and we aim to have the first deliverables ready by May 15, 2026."
    )
    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(12)
    run = body2.add_run(body2_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    body3_text = (
        "Please prepare a detailed project timeline and resource plan at your earliest "
        "convenience. We would appreciate receiving this by April 10, 2026, so that our "
        "internal stakeholders can review and approve the final scope before kickoff."
    )
    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(12)
    run = body3.add_run(body3_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    body4_text = (
        "Should you have any questions or require additional information about our brand "
        "guidelines or target demographics, please do not hesitate to reach out. We look "
        "forward to a productive collaboration."
    )
    body4 = doc.add_paragraph()
    body4.paragraph_format.space_after = Pt(12)
    run = body4.add_run(body4_text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(24)
    run = closing.add_run("Sincerely,")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    # --- Signature ---
    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_after = Pt(0)
    run = sig_name.add_run("James R. Mitchell")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    run = sig_title.add_run("Vice President, Business Development")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    sig_company = doc.add_paragraph()
    sig_company.paragraph_format.space_after = Pt(0)
    run = sig_company.add_run("Horizon Creative Solutions")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
