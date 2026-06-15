"""
Initial Setup: Add page numbers to the bottom center of every page in a user manual.
Task ID: writer_tech_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Ensure NO footer / page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    # --- Page 1: Title Page ---
    title = doc.add_heading('CloudSync Pro - User Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Version 3.2 | March 2026')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    company = doc.add_paragraph()
    company.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = company.add_run('NovaTech Solutions Inc.')
    run.font.size = Pt(16)
    run.bold = True

    addr = doc.add_paragraph()
    addr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = addr.add_run('1200 Innovation Boulevard, Suite 400\nSan Francisco, CA 94107')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # --- Page 2: Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1. Introduction', '3'),
        ('2. System Requirements', '3'),
        ('3. Installation Guide', '4'),
        ('4. Getting Started', '5'),
        ('5. Core Features', '5'),
        ('6. Advanced Configuration', '6'),
        ('7. Troubleshooting', '7'),
        ('8. Frequently Asked Questions', '8'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item} {"." * (50 - len(item))} {page}')
        run.font.size = Pt(11)

    doc.add_page_break()

    # --- Page 3: Introduction & System Requirements ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'CloudSync Pro is a comprehensive cloud storage and synchronization platform designed '
        'for enterprise teams that need reliable, secure file management across distributed '
        'workforces. Built on a microservices architecture, CloudSync Pro delivers sub-second '
        'sync latency with end-to-end AES-256 encryption for all data in transit and at rest.'
    )
    doc.add_paragraph(
        'Since its initial release in 2021, CloudSync Pro has been adopted by over 4,500 '
        'organizations worldwide, managing more than 12 petabytes of synchronized data. The '
        'platform integrates seamlessly with existing identity providers, supports granular '
        'access controls, and provides detailed audit logging for compliance requirements.'
    )
    doc.add_paragraph(
        'This user manual covers installation, configuration, daily usage workflows, and '
        'troubleshooting for CloudSync Pro version 3.2. For API documentation and developer '
        'resources, please refer to the separate Developer Reference Guide available at '
        'docs.novatech.io/cloudsync/api.'
    )

    doc.add_heading('2. System Requirements', level=1)
    doc.add_paragraph(
        'Before installing CloudSync Pro, ensure your system meets the following minimum '
        'requirements. Performance may vary on systems that meet only the minimum specifications; '
        'recommended specifications are provided for optimal experience.'
    )

    # Requirements table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Minimum', 'Recommended']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    reqs = [
        ['Operating System', 'Windows 10 / macOS 12 / Ubuntu 20.04', 'Windows 11 / macOS 14 / Ubuntu 22.04'],
        ['Processor', 'Intel i5-8250U / AMD Ryzen 5 2500U', 'Intel i7-12700 / AMD Ryzen 7 5800X'],
        ['Memory (RAM)', '4 GB', '16 GB'],
        ['Storage', '500 MB free space + sync data', '2 GB SSD + sync data'],
        ['Network', '10 Mbps broadband', '100 Mbps or higher'],
    ]
    for r, row_data in enumerate(reqs, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_page_break()

    # --- Page 4: Installation Guide ---
    doc.add_heading('3. Installation Guide', level=1)
    doc.add_heading('3.1 Windows Installation', level=2)
    doc.add_paragraph(
        'Download the latest installer from https://download.novatech.io/cloudsync/windows. '
        'Run the CloudSyncPro-3.2-Setup.exe file and follow the installation wizard. The '
        'default installation path is C:\\Program Files\\NovaTech\\CloudSync Pro. You may change '
        'this during installation if needed.'
    )
    steps_win = [
        'Double-click the downloaded installer file to launch the setup wizard.',
        'Accept the End User License Agreement after reviewing the terms.',
        'Choose the installation directory or accept the default path.',
        'Select optional components: Desktop shortcut, context menu integration, startup launch.',
        'Click Install and wait for the process to complete (typically 2-3 minutes).',
        'Launch CloudSync Pro from the Start Menu or desktop shortcut.',
    ]
    for step in steps_win:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('3.2 macOS Installation', level=2)
    doc.add_paragraph(
        'Download the DMG file from https://download.novatech.io/cloudsync/macos. Open the '
        'disk image and drag CloudSync Pro.app into your Applications folder. On first launch, '
        'macOS may prompt you to allow the application in System Preferences > Privacy & Security. '
        'Grant Full Disk Access for optimal synchronization performance.'
    )

    doc.add_heading('3.3 Linux Installation', level=2)
    doc.add_paragraph(
        'For Debian-based distributions, add the NovaTech repository and install via apt:'
    )
    code_block = doc.add_paragraph()
    run = code_block.add_run(
        'sudo add-apt-repository ppa:novatech/cloudsync\n'
        'sudo apt update\n'
        'sudo apt install cloudsync-pro'
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_page_break()

    # --- Page 5: Getting Started & Core Features ---
    doc.add_heading('4. Getting Started', level=1)
    doc.add_paragraph(
        'After installation, launch CloudSync Pro and sign in with your organization credentials. '
        'If your organization uses Single Sign-On (SSO), click "Sign in with SSO" and enter your '
        'company domain. The application will redirect you to your identity provider for '
        'authentication.'
    )
    doc.add_paragraph(
        'Once authenticated, the Setup Wizard will guide you through initial configuration:'
    )
    setup_items = [
        'Select the local folder to use as your sync root (default: ~/CloudSync).',
        'Choose which team folders to sync locally. Large teams may want selective sync.',
        'Configure bandwidth limits if you are on a metered or shared connection.',
        'Enable or disable desktop notifications for sync events and file changes.',
        'Review conflict resolution preferences (Keep Both, Keep Newest, or Ask Each Time).',
    ]
    for item in setup_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5. Core Features', level=1)
    doc.add_heading('5.1 Real-Time Synchronization', level=2)
    doc.add_paragraph(
        'CloudSync Pro monitors your designated sync folders using native filesystem event APIs '
        '(inotify on Linux, FSEvents on macOS, ReadDirectoryChangesW on Windows). When a file '
        'change is detected, the delta is computed using a rolling checksum algorithm similar to '
        'rsync, and only the modified blocks are uploaded. This approach reduces bandwidth usage '
        'by up to 90% compared to full-file uploads.'
    )

    doc.add_heading('5.2 Version History', level=2)
    doc.add_paragraph(
        'Every file synchronized through CloudSync Pro retains a complete version history for '
        '90 days (configurable up to 365 days on Enterprise plans). You can view, compare, and '
        'restore any previous version directly from the desktop application or the web dashboard. '
        'Version metadata includes the author, timestamp, file size, and a brief change summary '
        'when available.'
    )

    doc.add_page_break()

    # --- Page 6: Advanced Configuration ---
    doc.add_heading('6. Advanced Configuration', level=1)
    doc.add_heading('6.1 Configuration File', level=2)
    doc.add_paragraph(
        'Advanced users can modify the CloudSync Pro configuration file directly. The file is '
        'located at:'
    )
    paths = [
        'Windows: %APPDATA%\\NovaTech\\CloudSync\\config.yaml',
        'macOS: ~/Library/Application Support/NovaTech/CloudSync/config.yaml',
        'Linux: ~/.config/novatech/cloudsync/config.yaml',
    ]
    for p in paths:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_paragraph(
        'Key configuration parameters include:'
    )

    config_table = doc.add_table(rows=6, cols=3)
    config_table.style = 'Table Grid'
    config_headers = ['Parameter', 'Default', 'Description']
    for i, h in enumerate(config_headers):
        cell = config_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    config_data = [
        ['sync_interval_ms', '500', 'Polling interval in milliseconds when events are missed'],
        ['max_upload_kbps', '0 (unlimited)', 'Maximum upload bandwidth in kilobits per second'],
        ['max_download_kbps', '0 (unlimited)', 'Maximum download bandwidth in kilobits per second'],
        ['conflict_strategy', 'keep_both', 'Options: keep_both, keep_newest, keep_local, ask'],
        ['log_level', 'info', 'Options: debug, info, warn, error'],
    ]
    for r, row_data in enumerate(config_data, 1):
        for c, val in enumerate(row_data):
            config_table.cell(r, c).text = val

    doc.add_heading('6.2 Proxy Configuration', level=2)
    doc.add_paragraph(
        'If your network requires a proxy for outbound HTTPS connections, configure the proxy '
        'settings in the configuration file or via the Settings > Network panel in the application. '
        'CloudSync Pro supports HTTP, HTTPS, and SOCKS5 proxies. For NTLM-authenticated proxies, '
        'provide the domain\\username format in the proxy credentials field.'
    )

    doc.add_page_break()

    # --- Page 7: Troubleshooting ---
    doc.add_heading('7. Troubleshooting', level=1)
    doc.add_paragraph(
        'This section covers common issues encountered during installation and daily use of '
        'CloudSync Pro. For issues not listed here, contact NovaTech Support at '
        'support@novatech.io or visit the community forum at community.novatech.io.'
    )

    doc.add_heading('7.1 Sync Not Starting', level=2)
    doc.add_paragraph(
        'If synchronization does not begin after signing in, verify the following:'
    )
    troubleshoot_items = [
        'Confirm that your system clock is accurate. CloudSync Pro requires time synchronization '
        'within 60 seconds of NTP servers. Run "timedatectl status" on Linux or check Date & Time '
        'settings on Windows/macOS.',
        'Check that the sync root folder exists and has appropriate read/write permissions. On '
        'Linux, ensure your user owns the directory: chown -R $USER:$USER ~/CloudSync.',
        'Verify network connectivity to api.cloudsync.novatech.io on port 443. Use "curl -I '
        'https://api.cloudsync.novatech.io/health" to test the connection.',
        'Review the application logs for specific error messages. Logs are stored in the logs/ '
        'subdirectory of the configuration folder.',
    ]
    for item in troubleshoot_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 High CPU Usage', level=2)
    doc.add_paragraph(
        'Elevated CPU usage may occur during the initial sync when CloudSync Pro indexes all '
        'files in the sync root. This is normal and will subside once indexing completes. If '
        'high CPU usage persists, check for circular symbolic links in the sync directory tree '
        'and ensure the ignore_patterns configuration excludes large binary files (e.g., virtual '
        'machine images, database files) that change frequently.'
    )

    doc.add_heading('7.3 Conflict Resolution', level=2)
    doc.add_paragraph(
        'When two users modify the same file simultaneously, CloudSync Pro creates a conflict '
        'copy named "filename (Conflict - username - date).ext". To resolve conflicts manually, '
        'open both versions, merge the changes, save the canonical version, and delete the '
        'conflict copy. To change the default conflict behavior, update the conflict_strategy '
        'parameter in the configuration file.'
    )

    doc.add_page_break()

    # --- Page 8: FAQ ---
    doc.add_heading('8. Frequently Asked Questions', level=1)

    faqs = [
        ('Can I sync folders outside the designated sync root?',
         'Yes. Use the "Linked Folders" feature under Settings > Sync to add external folders. '
         'Each linked folder will appear as a separate top-level directory in your CloudSync '
         'workspace. Note that linked folders count toward your storage quota.'),
        ('How do I share a file with someone outside my organization?',
         'Right-click the file in your sync folder and select "Share via CloudSync". Choose '
         '"Anyone with the link" and set an expiration date and optional password. The recipient '
         'will receive a secure download link that does not require a CloudSync Pro account.'),
        ('What happens if I delete a file from the sync folder?',
         'Deleted files are moved to the CloudSync Trash, which is accessible from the web '
         'dashboard. Files remain in Trash for 30 days before permanent deletion. Organization '
         'administrators can extend this retention period up to 180 days.'),
        ('Does CloudSync Pro support end-to-end encryption?',
         'Yes. Enable Zero-Knowledge Encryption in Settings > Security. When enabled, files are '
         'encrypted locally before upload using a key derived from your passphrase. NovaTech '
         'servers never have access to unencrypted content. Note: this disables server-side '
         'search and file preview features.'),
        ('How do I migrate from another cloud storage provider?',
         'CloudSync Pro includes a Migration Assistant under Tools > Import. The assistant '
         'supports direct migration from Google Drive, Dropbox, OneDrive, and Box. During '
         'migration, existing folder structures and sharing permissions are preserved where '
         'possible. Large migrations (over 100 GB) should be scheduled during off-peak hours.'),
        ('What is the maximum file size supported?',
         'CloudSync Pro supports individual files up to 50 GB. Files larger than 100 MB are '
         'automatically uploaded using multipart upload with resumable transfer. If an upload '
         'is interrupted, it will resume from the last completed part on the next sync cycle.'),
    ]

    for q, a in faqs:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f'Q: {q}')
        q_run.bold = True
        doc.add_paragraph(f'A: {a}')
        doc.add_paragraph()  # spacing

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
