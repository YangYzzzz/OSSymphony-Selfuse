"""
Initial Setup: Academic paper with bibliography containing duplicate reference entries
Task ID: osworld_writer_duplicate_line_removal_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_duplicate_line_removal_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run(
        "Deep Learning Approaches for Natural Language Processing: "
        "A Comprehensive Survey"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()

    # Authors
    authors_para = doc.add_paragraph()
    authors_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors_run = authors_para.add_run(
        "Elena Vasquez\u00b9, James Thornton\u00b2, Priya Nair\u00b3"
    )
    authors_run.font.size = Pt(12)

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil_run = affil_para.add_run(
        "\u00b9Department of Computer Science, Stanford University\n"
        "\u00b2School of Engineering, MIT\n"
        "\u00b3Department of AI Research, Carnegie Mellon University"
    )
    affil_run.font.size = Pt(10)
    affil_run.italic = True

    doc.add_paragraph()

    # Abstract heading
    abstract_heading = doc.add_paragraph()
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    abstract_run = abstract_heading.add_run("Abstract")
    abstract_run.bold = True
    abstract_run.font.size = Pt(12)

    # Abstract body
    abstract_text = (
        "This paper presents a comprehensive survey of deep learning techniques "
        "applied to natural language processing (NLP) tasks. We review the evolution "
        "from recurrent neural networks (RNNs) to transformer-based architectures, "
        "examining their applications in machine translation, sentiment analysis, "
        "question answering, and text summarization. Our analysis covers 147 papers "
        "published between 2018 and 2024, identifying key trends, performance "
        "benchmarks, and open challenges. We conclude with recommendations for "
        "future research directions in low-resource languages, interpretability, "
        "and efficient inference."
    )
    abstract_para = doc.add_paragraph()
    abstract_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abstract_para.add_run(abstract_text).font.size = Pt(10)

    doc.add_paragraph()

    # Keywords
    keywords_para = doc.add_paragraph()
    kw_run = keywords_para.add_run("Keywords: ")
    kw_run.bold = True
    kw_run.font.size = Pt(10)
    keywords_para.add_run(
        "deep learning, natural language processing, transformers, BERT, GPT, survey"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # Section 1
    h1 = doc.add_paragraph()
    h1.add_run("1. Introduction").bold = True
    h1.runs[0].font.size = Pt(13)

    intro_text = (
        "Natural language processing has experienced a paradigm shift over the past "
        "decade with the advent of deep learning. Early NLP systems relied heavily "
        "on hand-crafted features and rule-based approaches, limiting their scalability "
        "and generalization capabilities. The introduction of word embeddings by "
        "Mikolov et al. [1] marked the beginning of a new era in which distributed "
        "representations captured semantic relationships between words.\n\n"
        "The development of sequence-to-sequence models [2] further advanced "
        "capabilities in machine translation and text generation. Attention mechanisms "
        "proposed by Bahdanau et al. [3] addressed the fixed-length bottleneck in "
        "encoder-decoder architectures, enabling models to focus selectively on "
        "relevant parts of the input. This culminated in the transformer architecture "
        "introduced by Vaswani et al. [4], which has since become the dominant "
        "framework for NLP research.\n\n"
        "Pre-trained language models such as BERT [5] and GPT [6] demonstrated "
        "that large-scale unsupervised pre-training followed by task-specific "
        "fine-tuning could achieve state-of-the-art results across diverse NLP "
        "benchmarks. Subsequent scaling laws research [7] revealed predictable "
        "improvements in model capabilities as a function of compute and data."
    )
    intro_para = doc.add_paragraph()
    intro_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_para.add_run(intro_text).font.size = Pt(11)

    doc.add_paragraph()

    # Section 2
    h2 = doc.add_paragraph()
    h2.add_run("2. Transformer-Based Architectures").bold = True
    h2.runs[0].font.size = Pt(13)

    transformer_text = (
        "The transformer architecture [4] relies entirely on self-attention "
        "mechanisms, dispensing with recurrence and convolutions. This design "
        "allows for greater parallelization during training and has proven "
        "highly effective at capturing long-range dependencies in text.\n\n"
        "BERT (Bidirectional Encoder Representations from Transformers) [5] "
        "introduced bidirectional pre-training using masked language modeling "
        "and next sentence prediction objectives. Fine-tuned BERT models achieved "
        "new state-of-the-art results on eleven NLP tasks at the time of publication. "
        "Subsequent variants including RoBERTa [8] optimized pre-training procedures "
        "through more data, longer training, and removal of the next sentence "
        "prediction objective.\n\n"
        "Autoregressive language models exemplified by GPT-2 [9] and GPT-3 [6] "
        "demonstrated impressive few-shot and zero-shot capabilities, challenging "
        "assumptions about the necessity of task-specific fine-tuning. The scaling "
        "behavior of these models was systematically studied by Kaplan et al. [7], "
        "who derived empirical scaling laws relating loss to model size, dataset "
        "size, and compute budget."
    )
    transformer_para = doc.add_paragraph()
    transformer_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    transformer_para.add_run(transformer_text).font.size = Pt(11)

    doc.add_paragraph()

    # Section 3
    h3 = doc.add_paragraph()
    h3.add_run("3. Applications in NLP Tasks").bold = True
    h3.runs[0].font.size = Pt(13)

    apps_text = (
        "Deep learning models have been successfully applied to a wide range of "
        "NLP tasks. In machine translation, the transformer architecture achieved "
        "substantial improvements over previous statistical methods [2]. Sentiment "
        "analysis benefited from contextualized representations, with BERT-based "
        "models substantially outperforming feature-engineering approaches.\n\n"
        "Question answering systems trained on large reading comprehension datasets "
        "approached human-level performance on benchmarks such as SQuAD. Named "
        "entity recognition and relation extraction tasks also saw significant "
        "gains from pre-trained representations. Text summarization, particularly "
        "abstractive approaches, improved substantially with sequence-to-sequence "
        "transformers.\n\n"
        "More recent work has explored multimodal NLP, combining text with visual "
        "information for tasks such as image captioning and visual question "
        "answering. Cross-lingual transfer learning enables models trained on "
        "high-resource languages to generalize to low-resource settings."
    )
    apps_para = doc.add_paragraph()
    apps_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    apps_para.add_run(apps_text).font.size = Pt(11)

    doc.add_paragraph()

    # Section 4
    h4 = doc.add_paragraph()
    h4.add_run("4. Conclusions").bold = True
    h4.runs[0].font.size = Pt(13)

    conclusion_text = (
        "This survey has examined the progression of deep learning methods in "
        "natural language processing from foundational word embeddings to large-scale "
        "pre-trained language models. The transformer architecture has emerged as "
        "the unifying framework, enabling breakthroughs across diverse NLP tasks.\n\n"
        "Persistent challenges remain in areas such as interpretability, computational "
        "efficiency, and performance on low-resource languages. Future research "
        "should explore parameter-efficient fine-tuning methods, robust evaluation "
        "protocols, and approaches to reduce environmental impact of large-scale "
        "model training. The integration of external knowledge sources and reasoning "
        "capabilities represents another promising avenue."
    )
    conclusion_para = doc.add_paragraph()
    conclusion_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    conclusion_para.add_run(conclusion_text).font.size = Pt(11)

    doc.add_paragraph()

    # Bibliography heading
    bib_heading = doc.add_paragraph()
    bib_heading.add_run("References").bold = True
    bib_heading.runs[0].font.size = Pt(13)

    # The 9 unique reference entries
    references = [
        "[1] Mikolov, T., Chen, K., Corrado, G., & Dean, J. (2013). Efficient estimation "
        "of word representations in vector space. arXiv preprint arXiv:1301.3781.",

        "[2] Sutskever, I., Vinyals, O., & Le, Q. V. (2014). Sequence to sequence learning "
        "with neural networks. Advances in Neural Information Processing Systems, 27.",

        "[3] Bahdanau, D., Cho, K., & Bengio, Y. (2015). Neural machine translation by "
        "jointly learning to align and translate. International Conference on Learning "
        "Representations (ICLR).",

        "[4] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., "
        "Kaiser, L., & Polosukhin, I. (2017). Attention is all you need. Advances in "
        "Neural Information Processing Systems, 30.",

        "[5] Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training "
        "of deep bidirectional transformers for language understanding. Proceedings of "
        "NAACL-HLT 2019, 4171\u20134186.",

        "[6] Brown, T. B., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., "
        "Neelakantan, A., Shyam, P., Sastry, G., Askell, A., et al. (2020). Language "
        "models are few-shot learners. Advances in Neural Information Processing Systems, 33.",

        "[7] Kaplan, J., McCandlish, S., Henighan, T., Brown, T. B., Chess, B., Child, R., "
        "Gray, S., Radford, A., Wu, J., & Amodei, D. (2020). Scaling laws for neural "
        "language models. arXiv preprint arXiv:2001.08361.",

        "[8] Liu, Y., Ott, M., Goyal, N., Du, J., Joshi, M., Chen, D., Levy, O., Lewis, M., "
        "Zettlemoyer, L., & Stoyanov, V. (2019). RoBERTa: A robustly optimized BERT "
        "pretraining approach. arXiv preprint arXiv:1907.11692.",

        "[9] Radford, A., Wu, J., Child, R., Luan, D., Amodei, D., & Sutskever, I. (2019). "
        "Language models are unsupervised multitask learners. OpenAI Blog, 1(8).",
    ]

    # The 5 duplicate entries (exact copies of some of the unique ones)
    # Duplicates of entries 1, 3, 5, 6, 8 (0-indexed: 0, 2, 4, 5, 7)
    duplicates = [
        references[0],  # duplicate of [1]
        references[2],  # duplicate of [3]
        references[4],  # duplicate of [5]
        references[5],  # duplicate of [6]
        references[7],  # duplicate of [8]
    ]

    # All 14 entries: 9 unique + 5 duplicates appended
    all_entries = references + duplicates

    for ref_text in all_entries:
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)
        ref_para.paragraph_format.space_after = Pt(4)
        ref_para.add_run(ref_text).font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
