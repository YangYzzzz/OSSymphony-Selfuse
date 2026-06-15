"""
Initial Setup: Employment contract with three key legal phrases, no footnotes.
Task ID: writer_hr_045
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Title --
    title = doc.add_heading('EMPLOYMENT AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -- Parties --
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('This Employment Agreement ("Agreement") is entered into as of March 15, 2025, '
                     'by and between Meridian Technologies, Inc., a Delaware corporation with its '
                     'principal offices located at 2400 Innovation Drive, Suite 800, San Jose, CA 95134 '
                     '("Employer"), and Rebecca Torres, an individual residing at 1847 Elm Street, '
                     'Apartment 12B, Palo Alto, CA 94301 ("Employee").')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 1: Position and Duties --
    doc.add_heading('1. Position and Duties', level=1)
    p = doc.add_paragraph()
    run = p.add_run('The Employer hereby employs the Employee as Senior Software Engineer, '
                     'reporting to the Vice President of Engineering, David Nakamura. The Employee '
                     'shall perform all duties and responsibilities customarily associated with this '
                     'position, including but not limited to: software architecture design, code review '
                     'and mentorship of junior developers, technical documentation, and participation '
                     'in cross-functional product planning meetings.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 2: Employment Type --
    doc.add_heading('2. Nature of Employment', level=1)
    p = doc.add_paragraph()
    run = p.add_run('This Agreement establishes an at-will employment relationship between the '
                     'Employer and the Employee. Either party may terminate this relationship at any '
                     'time, with or without cause, and with or without prior notice, subject to the '
                     'provisions outlined in Section 7 of this Agreement. Nothing in this Agreement '
                     'shall be construed to create a guarantee of continued employment for any '
                     'specific duration.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 3: Compensation --
    doc.add_heading('3. Compensation and Benefits', level=1)
    p = doc.add_paragraph()
    run = p.add_run('The Employee shall receive an annual base salary of $185,000.00, payable in '
                     'bi-weekly installments of $7,115.38 before applicable deductions. In addition, '
                     'the Employee shall be eligible for an annual performance bonus of up to 20% of '
                     'base salary, subject to achievement of individual and company performance targets '
                     'as determined by the Board of Directors. The Employee shall also receive a stock '
                     'option grant of 15,000 shares, vesting over a four-year period with a one-year cliff.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 4: Confidentiality --
    doc.add_heading('4. Confidentiality', level=1)
    p = doc.add_paragraph()
    run = p.add_run('The Employee acknowledges that during the course of employment, they will have '
                     'access to proprietary information, trade secrets, client lists, financial data, '
                     'and other confidential materials belonging to the Employer. The Employee agrees '
                     'to maintain strict confidentiality of all such information during and after the '
                     'term of employment, and shall not disclose any confidential information to third '
                     'parties without the prior written consent of the Employer.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 5: Non-Compete --
    doc.add_heading('5. Restrictive Covenants', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Upon termination of employment for any reason, the Employee agrees to abide by '
                     'a non-compete clause that restricts the Employee from engaging in any business '
                     'that directly competes with the Employer within a 50-mile radius of any Employer '
                     'office location for a period of twelve (12) months following the date of '
                     'termination. Additionally, the Employee shall not solicit any of the Employer\'s '
                     'clients, customers, or employees for a period of eighteen (18) months following '
                     'termination.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 6: Intellectual Property --
    doc.add_heading('6. Intellectual Property', level=1)
    p = doc.add_paragraph()
    run = p.add_run('All inventions, discoveries, designs, software code, algorithms, processes, and '
                     'works of authorship created by the Employee during the term of employment, whether '
                     'during working hours or using Employer resources, shall be the exclusive property '
                     'of the Employer. The Employee hereby assigns all rights, title, and interest in '
                     'such intellectual property to the Employer.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 7: Dispute Resolution --
    doc.add_heading('7. Dispute Resolution', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Any disputes arising out of or related to this Agreement shall be resolved through '
                     'an arbitration agreement between the parties. Both the Employer and the Employee '
                     'agree to submit any claims, controversies, or disputes to binding arbitration '
                     'administered by the American Arbitration Association in accordance with its '
                     'Employment Arbitration Rules. The arbitration shall take place in Santa Clara '
                     'County, California, and the decision of the arbitrator shall be final and binding '
                     'upon both parties.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Section 8: Governing Law --
    doc.add_heading('8. Governing Law', level=1)
    p = doc.add_paragraph()
    run = p.add_run('This Agreement shall be governed by and construed in accordance with the laws '
                     'of the State of California, without regard to its conflict of law principles. '
                     'Any provisions found to be unenforceable shall be severed from this Agreement, '
                     'and the remaining provisions shall continue in full force and effect.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # -- Signatures --
    doc.add_paragraph()  # blank line
    p = doc.add_paragraph()
    run = p.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date '
                     'first written above.')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('_________________________________')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Rebecca Torres, Employee')
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('_________________________________')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    run2 = p2.add_run('Jonathan Wei, CEO, Meridian Technologies, Inc.')
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
