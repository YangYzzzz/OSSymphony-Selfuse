"""
Initial Setup: Envelope template with address block as separate paragraphs
Task ID: writer_edit_069
Domain: libreoffice_writer

Creates an envelope/letter template with an address block where each line
is a separate paragraph (paragraph breaks between lines). The agent task
is to convert these to line breaks (soft returns) within a single paragraph.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_069'
# Task context says file is at ~/Desktop/
OUTPUT = f'{WORKDIR}/Desktop/envelope_template.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # --- Document settings ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Sender info (top of envelope/letter) ---
    sender_para = doc.add_paragraph()
    sender_run = sender_para.add_run('GlobalTech Solutions, Inc.')
    sender_run.bold = True
    sender_run.font.size = Pt(11)

    sender_addr = doc.add_paragraph('456 Commerce Boulevard, Floor 12')
    sender_addr.runs[0].font.size = Pt(11)

    city_para = doc.add_paragraph('Austin, TX 78701')
    city_para.runs[0].font.size = Pt(11)

    phone_para = doc.add_paragraph('Tel: (512) 555-0192')
    phone_para.runs[0].font.size = Pt(11)

    # --- Spacer ---
    doc.add_paragraph('')
    doc.add_paragraph('')

    # --- Date line ---
    date_para = doc.add_paragraph('March 15, 2025')
    date_para.runs[0].font.size = Pt(11)

    # --- Spacer ---
    doc.add_paragraph('')
    doc.add_paragraph('')

    # --- Address block (each line as its own paragraph — INITIAL STATE) ---
    # NOTE: These must remain as SEPARATE paragraphs for the initial state.
    # The agent's task is to merge them into one paragraph with line breaks.

    addr1 = doc.add_paragraph('Acme Corporation')
    addr1.runs[0].font.size = Pt(11)

    addr2 = doc.add_paragraph('1234 Innovation Drive')
    addr2.runs[0].font.size = Pt(11)

    addr3 = doc.add_paragraph('Suite 500')
    addr3.runs[0].font.size = Pt(11)

    addr4 = doc.add_paragraph('San Jose, CA 95134')
    addr4.runs[0].font.size = Pt(11)

    addr5 = doc.add_paragraph('United States')
    addr5.runs[0].font.size = Pt(11)

    # --- Spacer ---
    doc.add_paragraph('')
    doc.add_paragraph('')

    # --- Salutation ---
    salutation = doc.add_paragraph('Dear Procurement Team,')
    salutation.runs[0].font.size = Pt(11)

    # --- Spacer ---
    doc.add_paragraph('')

    # --- Body paragraphs ---
    body1 = doc.add_paragraph(
        'We are pleased to submit this formal proposal for enterprise software licensing '
        'and professional services as outlined in your Request for Proposal dated '
        'February 28, 2025. GlobalTech Solutions brings over fifteen years of experience '
        'delivering scalable technology solutions to Fortune 500 companies across North America.'
    )
    body1.runs[0].font.size = Pt(11)
    body1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('')

    body2 = doc.add_paragraph(
        'Our comprehensive platform integrates seamlessly with your existing infrastructure, '
        'providing real-time analytics, automated workflow management, and enterprise-grade '
        'security compliance. The enclosed documentation details our implementation timeline, '
        'support structure, and competitive pricing for a three-year partnership.'
    )
    body2.runs[0].font.size = Pt(11)
    body2.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('')

    body3 = doc.add_paragraph(
        'We welcome the opportunity to present our solution in greater detail at your '
        'convenience. Please do not hesitate to contact our account manager, Jennifer Walsh, '
        'at jennifer.walsh@globaltech.example.com or (512) 555-0198.'
    )
    body3.runs[0].font.size = Pt(11)
    body3.paragraph_format.space_after = Pt(6)

    # --- Spacer ---
    doc.add_paragraph('')
    doc.add_paragraph('')

    # --- Closing ---
    closing = doc.add_paragraph('Sincerely,')
    closing.runs[0].font.size = Pt(11)

    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    sig_name = doc.add_paragraph()
    sig_run = sig_name.add_run('Robert Chambers')
    sig_run.bold = True
    sig_run.font.size = Pt(11)

    sig_title = doc.add_paragraph('Vice President, Enterprise Sales')
    sig_title.runs[0].font.size = Pt(11)

    sig_company = doc.add_paragraph('GlobalTech Solutions, Inc.')
    sig_company.runs[0].font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # --- GUI-ready startup ---
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
