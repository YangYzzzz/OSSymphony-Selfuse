"""
Reward Script: Find all paragraph breaks in the address block and replace with line breaks (soft returns)
Task ID: writer_edit_069
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All 5 address lines exist within a SINGLE paragraph (merged address block)
  Component 2 (0.3 pts): That single paragraph contains exactly 4 soft returns (w:br without type)
  Component 3 (0.2 pts): Document has 4 fewer paragraphs than the original 32 (now 28),
                           confirming the 5-paragraph address block was collapsed into 1
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_069'

# Expected address lines (from task context)
EXPECTED_ADDRESS_LINES = [
    'Acme Corporation',
    '1234 Innovation Drive',
    'Suite 500',
    'San Jose, CA 95134',
    'United States',
]
# Initial file had 32 paragraphs; after merging 5 into 1, expect 28
EXPECTED_PARA_COUNT = 28
# Number of soft returns needed between 5 address lines
EXPECTED_SOFT_RETURNS = 4


def count_soft_returns_in_para(para):
    """Count line breaks (soft returns) in a paragraph.
    A soft return is <w:br/> without a w:type attribute (or not page/column/textWrapping).
    """
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    count = 0
    for run in para.runs:
        for br in run.element.findall('.//w:br', ns):
            br_type = br.attrib.get(qn('w:type'), '')
            # Soft return: no type attribute (empty string) — NOT page, column, or textWrapping
            if br_type not in ('page', 'column', 'textWrapping'):
                count += 1
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: The address block (5 lines: Acme Corporation, 1234 Innovation Drive, Suite 500,
    San Jose CA 95134, United States) must be merged from 5 separate paragraphs into
    a single paragraph with 4 soft returns (Shift+Enter line breaks) between them.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 5 address lines appear within a SINGLE paragraph (0.5 points)
    # In the initial file they are 5 separate paragraphs; in golden they are 1 paragraph.
    # We search for a paragraph whose text contains ALL 5 expected address lines.
    try:
        address_paragraph = None
        for para in doc.paragraphs:
            para_text = para.text
            # Check if ALL address lines appear in this paragraph
            if all(line in para_text for line in EXPECTED_ADDRESS_LINES):
                address_paragraph = para
                break

        if address_paragraph is not None:
            print(f"PASS: Component 1 — All 5 address lines found in a single paragraph (0.5 pts)")
            print(f"  Paragraph text: {address_paragraph.text!r}")
            total_score += 0.5
        else:
            # Check how many address lines are still in separate paragraphs
            para_texts = [p.text.strip() for p in doc.paragraphs]
            found_separate = [line for line in EXPECTED_ADDRESS_LINES if line in para_texts]
            print(f"FAIL: Component 1 — Address lines not merged into one paragraph (0.5 pts)")
            print(f"  Lines still as separate paragraphs: {found_separate}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: The merged address paragraph has exactly 4 soft returns (0.3 points)
    # This confirms that soft returns (Shift+Enter) were used, not just text concatenation.
    try:
        if address_paragraph is not None:
            soft_return_count = count_soft_returns_in_para(address_paragraph)
            if soft_return_count == EXPECTED_SOFT_RETURNS:
                print(f"PASS: Component 2 — Merged paragraph has exactly {EXPECTED_SOFT_RETURNS} soft returns (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Expected {EXPECTED_SOFT_RETURNS} soft returns (w:br), "
                      f"found {soft_return_count} (0.3 pts)")
        else:
            print(f"FAIL: Component 2 — Cannot check soft returns; address block not merged (0.3 pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Document paragraph count reduced by 4 (from 32 to 28) (0.2 points)
    # The 5 address paragraphs collapsed into 1 means 4 fewer paragraphs overall.
    try:
        actual_para_count = len(doc.paragraphs)
        if actual_para_count == EXPECTED_PARA_COUNT:
            print(f"PASS: Component 3 — Document paragraph count is {actual_para_count} "
                  f"(expected {EXPECTED_PARA_COUNT}) (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Expected {EXPECTED_PARA_COUNT} paragraphs, "
                  f"found {actual_para_count} (0.2 pts)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on the VM
file_path = f'{WORKDIR}/Desktop/envelope_template.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
