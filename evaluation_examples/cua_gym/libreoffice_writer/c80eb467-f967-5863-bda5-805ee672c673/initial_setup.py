"""
Initial Setup: Page numbering document with title, TOC, and main content
Task ID: writer_fs_093
Domain: libreoffice_writer

Creates a 25-page document:
- Page 1: Title page (First Page style, no page number)
- Pages 2-4: Table of Contents (3 pages of TOC entries)
- Pages 5-25: Main content chapters
All pages after page 1 show Arabic numerals 2-25 (default, no special formatting).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_093'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def add_page_number_footer(section, show=True):
    """Add a page number field to the footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    if not show:
        # Empty footer for title page
        for p in footer.paragraphs:
            p.clear()
        return
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fp.clear()
    # PAGE field code
    r1 = fp.add_run()
    fldChar1 = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fldChar1)
    r2 = fp.add_run()
    instrText = r2._element.makeelement(qn('w:instrText'), {})
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    r2._element.append(instrText)
    r3 = fp.add_run()
    fldChar2 = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fldChar2)


def add_filler_paragraphs(doc, count, prefix=""):
    """Add filler paragraphs to fill up a page."""
    for i in range(count):
        p = doc.add_paragraph(f"{prefix}{i+1}. " + "This section provides detailed analysis of the quarterly performance metrics and key deliverables across all departments.")


def create_initial():
    doc = Document()

    # ---- SECTION 1: Title Page (page 1) ----
    section1 = doc.sections[0]
    section1.page_width = Inches(8.5)
    section1.page_height = Inches(11)
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)
    section1.left_margin = Inches(1)
    section1.right_margin = Inches(1)

    # No page number on first page (titlePg and pgNumType set post-content)
    add_page_number_footer(section1, show=False)

    # Title content
    doc.add_paragraph("")  # spacer
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Quarterly Performance Report")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = subtitle.add_run("Fiscal Year 2025 — Q3 Review")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph("")

    org = doc.add_paragraph()
    org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = org.add_run("Meridian Analytics Corporation")
    run3.font.size = Pt(16)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = date_p.add_run("October 15, 2025")
    run4.font.size = Pt(14)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    conf = doc.add_paragraph()
    conf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run5 = conf.add_run("CONFIDENTIAL — Internal Distribution Only")
    run5.font.size = Pt(10)
    run5.italic = True
    run5.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # ---- SECTION 2: TOC Pages (pages 2-4) ----
    # New section with continuous page numbering (Arabic, starting from 2)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section2 = doc.sections[1]
    section2.page_width = Inches(8.5)
    section2.page_height = Inches(11)
    section2.top_margin = Inches(1)
    section2.bottom_margin = Inches(1)
    section2.left_margin = Inches(1)
    section2.right_margin = Inches(1)
    add_page_number_footer(section2, show=True)

    # Page number type: Arabic, continuing from page 1 (so page 2)
    # pgNumType set post-content

    toc_title = doc.add_heading("Table of Contents", level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TOC entries spanning 3 pages worth of content
    toc_entries = [
        ("1.", "Executive Summary", "5"),
        ("2.", "Revenue Analysis", "7"),
        ("  2.1", "North American Markets", "7"),
        ("  2.2", "European Markets", "8"),
        ("  2.3", "Asia-Pacific Markets", "9"),
        ("3.", "Operational Efficiency", "10"),
        ("  3.1", "Manufacturing Output", "10"),
        ("  3.2", "Supply Chain Optimization", "11"),
        ("  3.3", "Quality Assurance Metrics", "12"),
        ("4.", "Human Resources Overview", "13"),
        ("  4.1", "Headcount and Hiring", "13"),
        ("  4.2", "Employee Satisfaction Survey", "14"),
        ("  4.3", "Training Programs", "15"),
        ("5.", "Technology and Innovation", "16"),
        ("  5.1", "R&D Investment Summary", "16"),
        ("  5.2", "Product Development Pipeline", "17"),
        ("  5.3", "IT Infrastructure Upgrades", "18"),
        ("6.", "Financial Statements", "19"),
        ("  6.1", "Income Statement", "19"),
        ("  6.2", "Balance Sheet", "20"),
        ("  6.3", "Cash Flow Statement", "21"),
        ("7.", "Risk Assessment", "22"),
        ("  7.1", "Market Risks", "22"),
        ("  7.2", "Operational Risks", "23"),
        ("  7.3", "Regulatory Compliance", "23"),
        ("8.", "Strategic Outlook", "24"),
        ("  8.1", "Short-term Goals (Q4 2025)", "24"),
        ("  8.2", "Medium-term Strategy (2026)", "25"),
        ("  8.3", "Long-term Vision (2027-2030)", "25"),
        ("9.", "Appendices", "26"),
        ("  9.1", "Detailed Financial Tables", "26"),
        ("  9.2", "Survey Methodology", "27"),
        ("  9.3", "Glossary of Terms", "27"),
        ("10.", "Acknowledgments", "28"),
    ]

    for num, title_text, page in toc_entries:
        p = doc.add_paragraph()
        run_num = p.add_run(f"{num} ")
        run_num.font.size = Pt(11)
        run_title = p.add_run(title_text)
        run_title.font.size = Pt(11)
        if not num.startswith(" "):
            run_title.bold = True
        # Dots and page number
        dots = "." * max(2, 60 - len(num) - len(title_text))
        run_dots = p.add_run(f" {dots} ")
        run_dots.font.size = Pt(11)
        run_dots.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run_pg = p.add_run(page)
        run_pg.font.size = Pt(11)

    # Add extra spacing / entries to fill 3 pages of TOC
    doc.add_paragraph("")
    notes_heading = doc.add_paragraph()
    rn = notes_heading.add_run("List of Figures")
    rn.bold = True
    rn.font.size = Pt(13)

    figures = [
        ("Figure 1:", "Revenue Trends by Region (2023-2025)", "8"),
        ("Figure 2:", "Market Share Distribution", "9"),
        ("Figure 3:", "Manufacturing Output vs Target", "11"),
        ("Figure 4:", "Employee Satisfaction Scores", "14"),
        ("Figure 5:", "R&D Budget Allocation", "17"),
        ("Figure 6:", "Quarterly Cash Flow Summary", "21"),
        ("Figure 7:", "Risk Heat Map", "22"),
    ]
    for num, title_text, page in figures:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num} ")
        r1.italic = True
        r1.font.size = Pt(10)
        r2 = p.add_run(title_text)
        r2.font.size = Pt(10)
        dots = "." * max(2, 55 - len(num) - len(title_text))
        r3 = p.add_run(f" {dots} ")
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r4 = p.add_run(page)
        r4.font.size = Pt(10)

    doc.add_paragraph("")
    tables_heading = doc.add_paragraph()
    rt = tables_heading.add_run("List of Tables")
    rt.bold = True
    rt.font.size = Pt(13)

    tables_list = [
        ("Table 1:", "Revenue by Region (Q3 2025)", "7"),
        ("Table 2:", "Operational KPIs", "10"),
        ("Table 3:", "Headcount by Department", "13"),
        ("Table 4:", "R&D Project Status", "17"),
        ("Table 5:", "Income Statement Summary", "19"),
        ("Table 6:", "Balance Sheet Overview", "20"),
        ("Table 7:", "Risk Matrix Scores", "23"),
        ("Table 8:", "Strategic Initiative Timelines", "24"),
    ]
    for num, title_text, page in tables_list:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num} ")
        r1.italic = True
        r1.font.size = Pt(10)
        r2 = p.add_run(title_text)
        r2.font.size = Pt(10)
        dots = "." * max(2, 55 - len(num) - len(title_text))
        r3 = p.add_run(f" {dots} ")
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r4 = p.add_run(page)
        r4.font.size = Pt(10)

    # ---- SECTION 3: Main Content (pages 5-25, 21 pages) ----
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section3 = doc.sections[2]
    section3.page_width = Inches(8.5)
    section3.page_height = Inches(11)
    section3.top_margin = Inches(1)
    section3.bottom_margin = Inches(1)
    section3.left_margin = Inches(1)
    section3.right_margin = Inches(1)
    add_page_number_footer(section3, show=True)

    # Arabic numbering continuing from previous (so page 5, 6, 7, ...)
    # pgNumType set post-content

    # Chapter content across 21 pages
    chapters = [
        ("1. Executive Summary",
         "The third quarter of fiscal year 2025 has demonstrated robust growth across multiple business segments. Revenue reached $142.3 million, representing a 12.4% year-over-year increase. Our North American operations continued to lead with $78.5 million in revenue, while European markets showed the strongest growth rate at 18.2%. Key operational improvements include a 7% reduction in manufacturing costs and a 15% improvement in supply chain delivery times. Employee headcount grew to 3,247, with 156 new hires across engineering and sales departments."),
        ("2. Revenue Analysis",
         "Total consolidated revenue for Q3 2025 was $142.3 million, surpassing our internal target of $135 million by 5.4%. This performance was driven by strong demand in the enterprise software segment and favorable currency exchange rates in our international markets. Gross margin improved to 64.2%, up from 61.8% in the prior quarter, reflecting better cost controls and a shift toward higher-margin product offerings."),
        ("2.1 North American Markets",
         "North American revenue totaled $78.5 million, representing 55.2% of total revenue. The United States contributed $65.3 million while Canadian operations added $13.2 million. Enterprise client renewals remained strong at 94.7%, and new logo acquisitions added $8.4 million in annual recurring revenue. The federal government vertical showed particular strength with three new contracts worth $12.1 million combined."),
        ("2.2 European Markets",
         "European operations generated $38.7 million in revenue, marking an 18.2% year-over-year increase. Germany remained our largest European market at $14.2 million, followed by the United Kingdom at $10.8 million and France at $7.3 million. The DACH region expansion strategy is yielding results with 23 new enterprise accounts opened during the quarter."),
        ("2.3 Asia-Pacific Markets",
         "Asia-Pacific revenue reached $25.1 million, a 9.7% increase from Q2. Japan contributed $12.4 million, Australia $6.8 million, and Singapore $3.2 million. Our partnership with Toshiba Digital Solutions expanded our reach in the Japanese manufacturing sector, adding 15 new customer implementations."),
        ("3. Operational Efficiency",
         "Operational metrics showed improvement across all key areas. Total operating expenses were $91.2 million, representing 64.1% of revenue compared to 67.3% in Q2. The efficiency gains were primarily driven by manufacturing automation initiatives and the consolidation of three regional distribution centers into a single hub in Memphis, Tennessee."),
        ("3.1 Manufacturing Output",
         "Manufacturing output increased by 11.3% quarter-over-quarter, reaching 47,200 units. Defect rates decreased from 2.3% to 1.7%, resulting in $1.2 million in quality cost savings. The new automated assembly line in our Guadalajara facility contributed 8,500 units, exceeding its first-quarter target by 15%."),
        ("3.2 Supply Chain Optimization",
         "Average lead time decreased from 18 days to 15.3 days through improved supplier relationships and inventory management practices. On-time delivery rates improved to 96.4%, up from 93.1% in Q2. We renegotiated contracts with three major component suppliers, achieving an average 8.5% cost reduction on raw materials."),
        ("3.3 Quality Assurance Metrics",
         "Customer satisfaction scores improved to 4.6 out of 5.0, the highest in company history. Product return rates decreased to 1.2% from 1.8%. We achieved ISO 9001:2015 recertification with zero non-conformances noted. The new AI-powered inspection system identified 340 potential defects that would have been missed by traditional methods."),
        ("4. Human Resources Overview",
         "Our workforce expanded to 3,247 full-time employees, a net increase of 89 from Q2. Voluntary turnover decreased to 8.2% annualized, below the industry average of 12.5%. Total compensation expense was $48.7 million, including $6.2 million in performance bonuses and $2.1 million in equity-based compensation."),
        ("4.1 Headcount and Hiring",
         "We completed 156 new hires during Q3: 72 in Engineering, 38 in Sales, 24 in Operations, and 22 in Corporate Functions. Average time-to-fill decreased from 45 days to 38 days. Our engineering hiring pipeline grew to 450 qualified candidates. We launched partnerships with Georgia Tech and University of Waterloo for our internship program, receiving 1,200 applications for 40 positions."),
        ("4.2 Employee Satisfaction Survey",
         "The annual employee satisfaction survey yielded an 87% participation rate. Overall satisfaction scored 4.2 out of 5.0, up from 3.9 in 2024. Top-rated areas included team collaboration (4.5), management quality (4.3), and work-life balance (4.1). Areas identified for improvement include career development pathways (3.6) and internal communication (3.7)."),
        ("4.3 Training Programs",
         "Training investment totaled $1.8 million in Q3. We launched the Meridian Leadership Academy for high-potential managers, enrolling 45 participants. Technical certification completions increased by 34%, with 128 employees earning new industry certifications. The mentorship program expanded to 200 active mentor-mentee pairs."),
        ("5. Technology and Innovation",
         "R&D investment for Q3 was $18.4 million, representing 12.9% of revenue. We filed 12 new patent applications and received approvals for 8 previously submitted patents. Our technology roadmap progresses on schedule with the next major platform release targeted for Q1 2026."),
        ("5.1 R&D Investment Summary",
         "R&D spending was allocated across four primary areas: Platform Development ($8.2M), AI/ML Capabilities ($4.6M), Security Enhancements ($3.1M), and Infrastructure Modernization ($2.5M). Headcount in R&D grew to 487 engineers and researchers, including 42 PhDs. Our research partnership with MIT continued to generate promising results in natural language processing."),
        ("5.2 Product Development Pipeline",
         "Three major features reached general availability during Q3: Advanced Analytics Dashboard, Multi-tenant Architecture Support, and Real-time Collaboration Tools. Beta testing began for the AI-powered Document Processing module with 50 enterprise customers participating. Customer feedback has been overwhelmingly positive with a Net Promoter Score of 72."),
        ("5.3 IT Infrastructure Upgrades",
         "Cloud migration reached 78% completion, up from 62% in Q2. Monthly infrastructure costs decreased by $340,000 through right-sizing and reserved instance purchases. System availability maintained at 99.97% across all production services. We deployed a new CDN configuration that reduced average page load times by 23%."),
        ("6. Financial Statements",
         "The following sections present our consolidated financial statements for Q3 2025. All figures are reported in accordance with US GAAP standards and have been reviewed by our external auditors at Deloitte & Touche LLP."),
        ("6.1 Income Statement",
         "Total revenue of $142.3M with COGS of $50.9M resulted in gross profit of $91.4M (64.2% margin). Operating expenses totaled $91.2M, comprising R&D ($18.4M), Sales & Marketing ($24.6M), and G&A ($12.3M). Operating income was $51.1M, yielding a 35.9% operating margin. Net income after tax was $38.3M, or $1.47 per diluted share."),
        ("6.2 Balance Sheet",
         "Total assets stood at $892.4M as of September 30, 2025. Cash and equivalents were $234.7M, accounts receivable $87.3M, and property/equipment $156.2M. Total liabilities were $312.5M, including $125M in long-term debt. Shareholders equity increased to $579.9M. The current ratio improved to 2.8x from 2.5x in Q2."),
        ("6.3 Cash Flow Statement",
         "Operating cash flow was $52.1M, driven by strong net income and favorable working capital changes. Capital expenditures totaled $14.3M, primarily for the Guadalajara facility expansion and data center upgrades. Free cash flow was $37.8M. We repurchased $15M in company shares and paid $8.2M in dividends during the quarter."),
    ]

    for i, (heading, content) in enumerate(chapters):
        if i > 0:
            # Add page break between chapters to ensure proper pagination
            pb = doc.add_paragraph()
            pb.paragraph_format.page_break_before = True

        h = doc.add_heading(heading, level=2 if heading[0].isdigit() and '.' in heading and heading.split('.')[0].isdigit() and len(heading.split('.')[0]) <= 2 else 2)

        # Add the main paragraph
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)

        # Add some filler to ensure each chapter takes about a page
        filler_texts = [
            "Performance indicators continue to trend positively across all measured dimensions, reflecting the effectiveness of strategic initiatives implemented in prior quarters.",
            "Detailed breakdowns of these metrics are available in the appendix section of this report, with quarter-over-quarter comparisons provided for reference.",
            "Management remains confident in sustaining this trajectory through continued investment in operational excellence and market expansion strategies.",
            "Cross-functional team reviews have confirmed alignment between departmental objectives and the overall corporate strategy outlined at the annual planning session.",
        ]
        for ft in filler_texts:
            fp = doc.add_paragraph(ft)
            fp.paragraph_format.space_after = Pt(4)

    # ---- Post-content: Set page numbering via direct XML on the correct sectPr elements ----
    ns_w = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body
    pPr_sectPrs = body.findall('.//w:pPr/w:sectPr', ns_w)
    body_sectPrs = body.findall('w:sectPr', ns_w)
    all_sectPrs = pPr_sectPrs + body_sectPrs

    for sp in all_sectPrs:
        for old in sp.findall('w:pgNumType', ns_w):
            sp.remove(old)
        for old in sp.findall('w:titlePg', ns_w):
            sp.remove(old)

    # Section 0: title page - add titlePg to suppress page number on page 1
    titlePg = all_sectPrs[0].makeelement(qn('w:titlePg'), {})
    all_sectPrs[0].append(titlePg)

    # Section 1 (TOC): Arabic decimal (default continuing numbering = pages 2-4)
    pgNum1 = all_sectPrs[1].makeelement(qn('w:pgNumType'), {
        qn('w:fmt'): 'decimal'
    })
    all_sectPrs[1].append(pgNum1)

    # Section 2 (main content): Arabic decimal (continuing = pages 5-25)
    pgNum2 = all_sectPrs[2].makeelement(qn('w:pgNumType'), {
        qn('w:fmt'): 'decimal'
    })
    all_sectPrs[2].append(pgNum2)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


create_initial()

# GUI-ready startup
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
