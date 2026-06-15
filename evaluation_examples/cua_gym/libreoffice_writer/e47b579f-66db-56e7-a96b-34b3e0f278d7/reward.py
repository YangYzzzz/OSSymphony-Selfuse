"""
Reward Script: Find and replace 'e.g.' with 'e.g.,' in academic paper
Task ID: writer_frd_034
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): No remaining 'e.g.' without trailing comma
  Component 2 (0.3): Total 'e.g.,' count is 16 (11 fixed + 5 original)
  Component 3 (0.2): No double commas introduced AND document structure preserved
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_034'


def persist_app_state(domain: str):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph texts for analysis
    all_texts = [para.text for para in doc.paragraphs]
    full_text = '\n'.join(all_texts)

    # Count e.g. without trailing comma: pattern matches 'e.g.' followed by non-comma
    # Also handle 'e.g.' at end of string (no following char)
    eg_no_comma = len(re.findall(r'e\.g\.(?!,)', full_text))
    # Count e.g. with trailing comma
    eg_with_comma = len(re.findall(r'e\.g\.,', full_text))
    # Count double commas (corruption check)
    double_commas = len(re.findall(r'e\.g\.,,', full_text))

    print(f"Analysis: e.g. without comma={eg_no_comma}, e.g. with comma={eg_with_comma}, double commas={double_commas}")

    # Component 1: No remaining 'e.g.' without trailing comma (0.5 points)
    # Initial has 11 such instances; golden should have 0
    try:
        if eg_no_comma == 0:
            print(f"PASS: Component 1 — All 'e.g.' instances have trailing commas (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Found {eg_no_comma} 'e.g.' without trailing comma (expected 0)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Total 'e.g.,' count is 16 (0.3 points)
    # Initial has 5, golden should have 16 (5 original + 11 fixed)
    # Award partial: if at least some were fixed but not all
    try:
        if eg_with_comma >= 16:
            print(f"PASS: Component 2 — Found {eg_with_comma} 'e.g.,' instances (expected 16) (0.3 pts)")
            total_score += 0.3
        elif eg_with_comma > 5:
            # Partial credit: some were fixed
            fixed_count = eg_with_comma - 5
            partial = 0.3 * (fixed_count / 11.0)
            print(f"PARTIAL: Component 2 — Found {eg_with_comma} 'e.g.,' instances ({fixed_count}/11 fixed) ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Found {eg_with_comma} 'e.g.,' instances (expected 16, baseline is 5)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Replacements done cleanly — no double commas AND all fixed (0.2 points)
    # This is anchored to the task change: only awards points if replacements happened
    # correctly (eg_no_comma == 0) AND no corruption was introduced (no double commas)
    try:
        num_paras = len(doc.paragraphs)
        no_double = double_commas == 0
        all_fixed = eg_no_comma == 0

        if no_double and all_fixed and num_paras == 34:
            print(f"PASS: Component 3 — Clean replacements, no double commas, structure preserved (0.2 pts)")
            total_score += 0.2
        else:
            if not all_fixed:
                print(f"FAIL: Component 3 — Replacements not complete ({eg_no_comma} unfixed)")
            if not no_double:
                print(f"FAIL: Component 3 — Found {double_commas} double comma instances ('e.g.,,')")
            if num_paras != 34:
                print(f"FAIL: Component 3 — Paragraph count is {num_paras}, expected 34")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved edits
persist_app_state("libreoffice_writer")

# Run verification
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
