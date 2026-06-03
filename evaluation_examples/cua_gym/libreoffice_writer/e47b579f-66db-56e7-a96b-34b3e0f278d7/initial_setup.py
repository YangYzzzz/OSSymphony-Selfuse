"""
Initial Setup: Academic paper with 'e.g.' instances needing comma correction
Task ID: writer_frd_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Computational Approaches to Natural Language Understanding: A Survey', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author info ---
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Dr. Elena Vasquez, Dr. Rajesh Krishnamurthy, Dr. Sarah Mitchell')
    run.font.size = Pt(11)
    run.font.italic = True
    author2 = doc.add_paragraph()
    author2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = author2.add_run('Department of Computer Science, Stanford University')
    run2.font.size = Pt(10)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    # Instance 1: "e.g." without comma
    doc.add_paragraph(
        'Natural language processing (NLP) has undergone significant transformations in recent years, '
        'particularly with the advent of deep learning methods. Various architectures have been proposed '
        'for core NLP tasks (e.g. machine translation, sentiment analysis, and named entity recognition). '
        'This survey examines the evolution of computational approaches from rule-based systems to '
        'modern transformer-based models, highlighting key milestones and remaining challenges.'
    )

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)
    # Instance 2: "e.g." without comma
    doc.add_paragraph(
        'The field of natural language understanding has expanded rapidly over the past decade. '
        'Early systems relied on hand-crafted rules and linguistic features (e.g. parse trees and '
        'morphological analysis) to process text. These systems, while interpretable, suffered from '
        'limited coverage and poor generalization to unseen data.'
    )
    # Already correct: "e.g.," with comma (should not be doubled)
    doc.add_paragraph(
        'Modern approaches leverage large-scale pre-trained models, such as BERT (Devlin et al., 2019) '
        'and GPT (Radford et al., 2018). These models have demonstrated impressive performance on '
        'standard benchmarks (e.g., GLUE, SuperGLUE, and SQuAD) and have become the de facto '
        'foundation for many downstream applications.'
    )
    # Instance 3: "e.g." without comma
    doc.add_paragraph(
        'Despite these advances, significant challenges remain. Low-resource languages (e.g. Swahili, '
        'Tagalog, and Welsh) still lack sufficient training data, and domain-specific applications '
        'require careful adaptation strategies.'
    )

    # --- 2. Background ---
    doc.add_heading('2. Background and Related Work', level=1)
    # Instance 4: "e.g." without comma
    doc.add_paragraph(
        'Statistical methods dominated NLP research throughout the 1990s and 2000s. Techniques such as '
        'Hidden Markov Models and Conditional Random Fields were applied to sequential labeling tasks '
        '(e.g. part-of-speech tagging and chunking). The introduction of word embeddings by Mikolov '
        'et al. (2013) marked a paradigm shift toward distributed representations.'
    )
    # Already correct: "e.g.," with comma
    doc.add_paragraph(
        'Transfer learning has become a cornerstone of modern NLP. Pre-trained language models '
        'capture general linguistic knowledge that can be fine-tuned for specific tasks (e.g., '
        'text classification, question answering, and summarization). This approach has dramatically '
        'reduced the need for task-specific architectural innovations.'
    )
    # Instance 5: "e.g." without comma
    doc.add_paragraph(
        'Several survey papers have examined subsets of this field. Previous reviews focused on '
        'specific application domains (e.g. biomedical text mining or legal document analysis) '
        'rather than providing a comprehensive cross-domain perspective.'
    )

    # --- 3. Methodology ---
    doc.add_heading('3. Methodology', level=1)
    # Instance 6: "e.g." without comma
    doc.add_paragraph(
        'Our survey methodology follows established systematic review protocols. We collected papers '
        'from major venues (e.g. ACL, EMNLP, NAACL, and NeurIPS) published between 2015 and 2025. '
        'We applied inclusion criteria requiring empirical evaluation on at least one benchmark dataset.'
    )
    # Instance 7: "e.g." without comma
    doc.add_paragraph(
        'For each paper, we extracted key metadata including the model architecture, training data '
        'characteristics, and evaluation metrics. We categorized approaches along several dimensions '
        '(e.g. supervised vs. self-supervised learning and monolingual vs. multilingual settings). '
        'This taxonomy enables systematic comparison across methods.'
    )

    # --- 4. Architectures ---
    doc.add_heading('4. Model Architectures', level=1)
    # Already correct: "e.g.," with comma
    doc.add_paragraph(
        'Transformer-based architectures dominate current research. Encoder-only models (e.g., '
        'BERT and RoBERTa) excel at understanding tasks, while decoder-only models (e.g., GPT-3 '
        'and PaLM) demonstrate strong generative capabilities.'
    )
    # Instance 8: "e.g." without comma
    doc.add_paragraph(
        'Encoder-decoder architectures (e.g. T5 and BART) provide flexibility for both understanding '
        'and generation tasks. These models have shown particular strength in tasks requiring '
        'conditional text generation, such as machine translation and abstractive summarization.'
    )
    # Instance 9: "e.g." without comma
    doc.add_paragraph(
        'Recent work has explored efficient alternatives to full attention mechanisms. Sparse '
        'attention patterns (e.g. local windows and strided attention) reduce the quadratic '
        'complexity of standard self-attention while maintaining competitive performance on '
        'long-document tasks.'
    )

    # --- 5. Evaluation ---
    doc.add_heading('5. Evaluation and Benchmarks', level=1)
    # Instance 10: "e.g." without comma
    doc.add_paragraph(
        'Evaluation practices in NLP have evolved alongside model capabilities. Traditional metrics '
        '(e.g. BLEU for translation and F1 for extraction) remain widely used but have known '
        'limitations. Human evaluation studies provide more nuanced assessments but are expensive '
        'and difficult to standardize.'
    )
    # Already correct: "e.g.," with comma
    doc.add_paragraph(
        'Multi-task benchmarks (e.g., GLUE and its successor SuperGLUE) have been instrumental in '
        'driving progress. However, recent work has questioned whether high benchmark scores truly '
        'reflect genuine language understanding or merely pattern matching.'
    )

    # --- 6. Discussion ---
    doc.add_heading('6. Discussion', level=1)
    # Instance 11: "e.g." without comma
    doc.add_paragraph(
        'Several open challenges deserve attention. Ethical considerations (e.g. bias in training '
        'data and privacy concerns) must be addressed as these models are deployed in high-stakes '
        'applications. Additionally, the environmental cost of training large language models has '
        'become a growing concern in the research community.'
    )
    doc.add_paragraph(
        'Future directions include developing more sample-efficient learning methods, improving '
        'model interpretability, and creating more robust evaluation frameworks that better capture '
        'the nuances of human language understanding.'
    )

    # --- 7. Conclusion ---
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        'This survey has provided a comprehensive overview of computational approaches to natural '
        'language understanding. We have traced the evolution from rule-based systems through '
        'statistical methods to modern neural architectures. While significant progress has been '
        'made, fundamental challenges in robustness, fairness, and genuine comprehension remain. '
        'We hope this work serves as a useful reference for researchers entering the field.'
    )

    # --- References ---
    doc.add_heading('References', level=1)
    refs = [
        'Devlin, J., Chang, M.W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. NAACL-HLT.',
        'Mikolov, T., Chen, K., Corrado, G., & Dean, J. (2013). Efficient estimation of word representations in vector space. ICLR Workshop.',
        'Radford, A., Narasimhan, K., Salimans, T., & Sutskever, I. (2018). Improving language understanding by generative pre-training. OpenAI Technical Report.',
        'Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention is all you need. NeurIPS.',
        'Wang, A., Singh, A., Michael, J., et al. (2018). GLUE: A multi-task benchmark and analysis platform for natural language understanding. ICLR.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
