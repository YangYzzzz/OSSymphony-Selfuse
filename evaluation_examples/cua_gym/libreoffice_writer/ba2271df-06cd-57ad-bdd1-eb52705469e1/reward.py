"""
Reward Script: Document Comparison — Accept text changes, reject formatting changes
Task ID: writer_fp_033
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.40): Rewritten paragraphs contain revised text content
  - Component 2 (0.20): Added sentences are present in extended paragraphs
  - Component 3 (0.25): Text is changed BUT heading font sizes remain original (compound)
  - Component 4 (0.15): Text is changed BUT paragraph spacing remains original (compound)

All components are anchored to the text change (fails on initial, passes on golden).
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_033'

# --- Ground truth from task context ---
# Text changes: paragraphs rewritten or extended in revised doc
# P16: completely rewritten (Executive Summary first para)
# P34: completely rewritten (Operational continuity para)
# P48: completely rewritten (ROI projections para)
# P30: sentence added at end (Phase 3 description)
# P54: sentence added at end (Customer satisfaction para)

# Unique substrings in REVISED text (absent from original)
REWRITTEN_MARKERS = {
    16: "reached end-of-life across multiple critical components",
    34: "meticulous capacity engineering",
    48: "break-even at month 30 rather than the originally projected month 34",
}

ADDED_SENTENCE_MARKERS = {
    30: "accessibility review will be conducted during this phase",
    54: "25% reduction in call center volume",
}

# Original heading font size (EMU): 203200 = 16pt
ORIGINAL_HEADING_SIZE = 203200
HEADING_INDICES = [5, 15, 20, 26, 32, 38, 45, 51, 57]

# Original paragraph spacing: space_after = 101600 EMU = 8pt
ORIGINAL_SPACE_AFTER = 101600
SPACING_CHECK_INDICES = [16, 17, 18, 21, 22, 23, 24, 27, 28, 29, 33, 34, 35, 46, 47, 48, 52, 53, 54, 58, 59, 60]


def persist_app_state(domain: str):
    """Try to save any unsaved changes in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_paras = len(doc.paragraphs)
    if num_paras < 60:
        print(f"CRITICAL: Document has only {num_paras} paragraphs, expected ~63")
        print("REWARD: 0.0")
        return 0.0

    # ---- Helper: check if ANY rewritten text is present ----
    # This is the "anchor" — if no text changes are present, the task hasn't been done
    any_text_changed = False

    # Component 1: Rewritten paragraphs contain revised text (0.40 points)
    # These paragraphs were completely rewritten in the revised doc.
    # FAILS on initial (original text), PASSES on golden (revised text).
    try:
        rewrite_passed = 0
        rewrite_total = len(REWRITTEN_MARKERS)

        for idx, marker in REWRITTEN_MARKERS.items():
            para_text = doc.paragraphs[idx].text
            if marker in para_text:
                print(f"PASS: P{idx} contains revised text: '{marker[:50]}...'")
                rewrite_passed += 1
            else:
                print(f"FAIL: P{idx} missing revised text. Found: '{para_text[:80]}...'")

        if rewrite_passed == rewrite_total:
            total_score += 0.40
            any_text_changed = True
            print(f"PASS: Component 1 — All {rewrite_total} rewritten paragraphs verified (0.40 pts)")
        elif rewrite_passed > 0:
            partial = round(0.40 * (rewrite_passed / rewrite_total), 2)
            total_score += partial
            any_text_changed = True
            print(f"PARTIAL: Component 1 — {rewrite_passed}/{rewrite_total} rewrites ({partial} pts)")
        else:
            print(f"FAIL: Component 1 — No rewritten paragraphs found (0 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Added sentences are present (0.20 points)
    # These paragraphs had new sentences appended in the revised doc.
    # FAILS on initial (no added text), PASSES on golden (text appended).
    try:
        added_passed = 0
        added_total = len(ADDED_SENTENCE_MARKERS)

        for idx, marker in ADDED_SENTENCE_MARKERS.items():
            para_text = doc.paragraphs[idx].text
            if marker in para_text:
                print(f"PASS: P{idx} contains added sentence: '{marker[:50]}...'")
                added_passed += 1
            else:
                print(f"FAIL: P{idx} missing added sentence. Found: '{para_text[:80]}...'")

        if added_passed == added_total:
            total_score += 0.20
            if not any_text_changed:
                any_text_changed = True
            print(f"PASS: Component 2 — All {added_total} added sentences verified (0.20 pts)")
        elif added_passed > 0:
            partial = round(0.20 * (added_passed / added_total), 2)
            total_score += partial
            any_text_changed = True
            print(f"PARTIAL: Component 2 — {added_passed}/{added_total} additions ({partial} pts)")
        else:
            print(f"FAIL: Component 2 — No added sentences found (0 pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text changed AND heading font sizes remain original (0.25 points)
    # COMPOUND CHECK: requires BOTH text changes (Component 1/2 anchor) AND original formatting.
    # FAILS on initial (no text changes even though formatting is original).
    # PASSES on golden (text changed + formatting kept original = formatting changes rejected).
    try:
        if not any_text_changed:
            print(f"FAIL: Component 3 — Text changes not present, cannot verify formatting rejection (0 pts)")
        else:
            headings_correct = 0
            headings_total = 0

            for idx in HEADING_INDICES:
                if idx < num_paras and doc.paragraphs[idx].runs:
                    headings_total += 1
                    font_size = doc.paragraphs[idx].runs[0].font.size
                    if font_size == ORIGINAL_HEADING_SIZE:
                        headings_correct += 1
                    else:
                        print(f"FAIL: P{idx} heading font size = {font_size}, expected {ORIGINAL_HEADING_SIZE}")

            if headings_total > 0 and headings_correct == headings_total:
                total_score += 0.25
                print(f"PASS: Component 3 — Text changed + all {headings_total} headings retain original size (0.25 pts)")
            elif headings_correct > 0:
                partial = round(0.25 * (headings_correct / headings_total), 2)
                total_score += partial
                print(f"PARTIAL: Component 3 — {headings_correct}/{headings_total} headings correct ({partial} pts)")
            else:
                print(f"FAIL: Component 3 — All headings have non-original font size (0 pts)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Text changed AND paragraph spacing remains original (0.15 points)
    # COMPOUND CHECK: requires text changes anchor + original spacing preserved.
    # FAILS on initial (no text changes). PASSES on golden (text + original spacing).
    try:
        if not any_text_changed:
            print(f"FAIL: Component 4 — Text changes not present, cannot verify spacing rejection (0 pts)")
        else:
            spacing_correct = 0
            spacing_total = 0

            for idx in SPACING_CHECK_INDICES:
                if idx < num_paras:
                    spacing_total += 1
                    space_after = doc.paragraphs[idx].paragraph_format.space_after
                    if space_after == ORIGINAL_SPACE_AFTER:
                        spacing_correct += 1
                    else:
                        print(f"FAIL: P{idx} space_after = {space_after}, expected {ORIGINAL_SPACE_AFTER}")

            if spacing_total > 0 and spacing_correct == spacing_total:
                total_score += 0.15
                print(f"PASS: Component 4 — Text changed + all {spacing_total} spacings retain original (0.15 pts)")
            elif spacing_correct > 0:
                partial = round(0.15 * (spacing_correct / spacing_total), 2)
                total_score += partial
                print(f"PARTIAL: Component 4 — {spacing_correct}/{spacing_total} spacings correct ({partial} pts)")
            else:
                print(f"FAIL: Component 4 — All paragraphs have non-original spacing (0 pts)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
