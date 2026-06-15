"""
Initial Setup: Document comparison task - original proposal and revised version
Task ID: writer_fp_033
Domain: libreoffice_writer

Creates:
  /home/user/writer_fp_033.docx — Original proposal document (8 pages)
  /home/user/Documents/proposal_revised.docx — Revised version with text+formatting changes
  Opens writer_fp_033.docx in LibreOffice Writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
REVISED = f'{WORKDIR}/Documents/proposal_revised.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# ── Shared content helpers ──────────────────────────────────────────────

HEADING_FONT_ORIGINAL = Pt(16)
HEADING_FONT_REVISED = Pt(20)  # formatting change in revised

BODY_SPACE_AFTER_ORIGINAL = Pt(8)
BODY_SPACE_AFTER_REVISED = Pt(14)  # formatting change in revised


def add_heading_styled(doc, text, font_size=HEADING_FONT_ORIGINAL):
    """Add a heading with explicit font size."""
    para = doc.add_heading(text, level=1)
    for run in para.runs:
        run.font.size = font_size
    return para


def add_body_para(doc, text, space_after=BODY_SPACE_AFTER_ORIGINAL):
    """Add a body paragraph with explicit spacing."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = space_after
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.name = "Liberation Serif"
    return para


# ── Content that stays the same in both versions ────────────────────────

TITLE = "Strategic Technology Modernization Proposal"
SUBTITLE = "Prepared for Meridian Financial Group"
DATE_LINE = "March 2026"

# Section headings (8 sections to fill ~8 pages)
SECTIONS = [
    "Executive Summary",
    "Current Infrastructure Assessment",
    "Proposed Migration Strategy",
    "Risk Analysis and Mitigation",
    "Implementation Timeline",
    "Budget and Resource Allocation",
    "Expected Outcomes and KPIs",
    "Appendix: Technical Specifications",
]

# Paragraphs per section (shared between original and revised except where noted)
SECTION_CONTENT_ORIGINAL = {
    "Executive Summary": [
        "Meridian Financial Group currently operates on a legacy technology stack that was originally deployed in 2014. Over the past decade, this infrastructure has served the organization well, but increasing operational costs, security vulnerabilities, and scalability limitations have made modernization a strategic priority for 2026 and beyond.",
        "This proposal outlines a comprehensive three-phase migration plan designed to transition Meridian's core banking platform, customer relationship management system, and data analytics infrastructure to a modern cloud-native architecture. The estimated total investment is $4.7 million over 24 months, with projected annual savings of $1.2 million beginning in Year 3.",
        "Our assessment indicates that delaying modernization beyond Q3 2026 would expose the organization to significant regulatory compliance risks, particularly with the upcoming DORA framework requirements. The proposed timeline accounts for minimal business disruption while maintaining full regulatory compliance throughout the transition period.",
    ],
    "Current Infrastructure Assessment": [
        "The existing technology landscape at Meridian comprises 47 on-premises servers distributed across two data centers in Chicago and Newark. The primary database cluster runs Oracle 12c with approximately 18 terabytes of structured financial data. Backup and disaster recovery systems utilize tape-based archival with a recovery time objective of 72 hours.",
        "Network infrastructure relies on a hub-and-spoke topology connecting 23 branch offices to the Chicago primary data center. Current bandwidth utilization peaks at 78% during end-of-month processing cycles, causing noticeable performance degradation for branch users accessing the core banking application during these periods.",
        "Security assessments conducted in January 2026 identified 14 critical vulnerabilities in the current stack, including outdated TLS configurations, unpatched middleware components, and insufficient network segmentation between the production and development environments. Six of these findings require remediation within 90 days per regulatory guidance.",
        "The customer-facing web portal, built on Java EE 7, experiences average response times of 3.2 seconds during peak hours. Mobile banking functionality is limited to basic account viewing, with no support for real-time payments, biometric authentication, or personalized financial insights.",
    ],
    "Proposed Migration Strategy": [
        "We recommend a phased approach that prioritizes risk reduction and quick wins in Phase 1, followed by core system migration in Phase 2, and advanced capability deployment in Phase 3. Each phase includes dedicated testing periods, rollback procedures, and stakeholder review gates.",
        "Phase 1 (Months 1-6) focuses on cloud foundation setup, including establishing a multi-region AWS environment with proper security controls, identity management integration with existing Active Directory, and network connectivity via AWS Direct Connect to both data center locations. During this phase, non-critical workloads such as the employee intranet, development environments, and reporting dashboards will be migrated first.",
        "Phase 2 (Months 7-16) addresses the core banking platform migration. This involves deploying containerized microservices on Amazon EKS, migrating the Oracle database to Amazon Aurora PostgreSQL, and implementing real-time event streaming via Apache Kafka for inter-service communication. A parallel-run period of 60 days is planned to validate data consistency.",
        "Phase 3 (Months 17-24) delivers advanced capabilities including an AI-powered fraud detection engine, a customer analytics platform leveraging Amazon SageMaker, and a completely redesigned mobile banking experience built on React Native with biometric authentication and real-time push notifications.",
    ],
    "Risk Analysis and Mitigation": [
        "Data migration accuracy represents the highest risk factor in this initiative. Financial records must maintain absolute integrity through the transition. Our mitigation strategy includes automated reconciliation scripts that compare source and target databases row-by-row, hash-based verification of account balances, and a dedicated data quality team conducting manual audits of randomly selected records.",
        "Operational continuity during the parallel-run period requires careful capacity planning. We have allocated 30% additional compute resources during Phase 2 to handle the overhead of running both legacy and modern systems simultaneously. Branch office connectivity will be maintained through dual routing paths.",
        "Regulatory compliance must be maintained throughout the migration. Our approach includes pre-migration compliance audits, continuous monitoring during transition, and post-migration certification. We have engaged Deloitte's financial services technology practice to provide independent compliance validation at each phase gate.",
        "Staff training and change management pose significant risks to user adoption. We propose a train-the-trainer model beginning 90 days before each phase go-live, supplemented by on-site support teams during the first 30 days of each transition. A dedicated help desk channel will be established for migration-related issues.",
    ],
    "Implementation Timeline": [
        "The implementation follows an aggressive but achievable 24-month schedule. Key milestones are structured around quarterly review points where the steering committee evaluates progress, budget adherence, and risk indicators before authorizing continuation to the next phase.",
        "Quarter 1 (January-March 2027): Cloud environment provisioning, security baseline configuration, Direct Connect establishment, and migration of development and staging environments. Deliverable: Fully operational cloud landing zone with passing security audits.",
        "Quarter 2 (April-June 2027): Employee intranet migration, reporting dashboard replatforming, and initial API gateway deployment. Begin core banking application containerization in development environment. Deliverable: First production workloads running in cloud.",
        "Quarter 3-4 (July-December 2027): Core banking microservices deployment, database migration, parallel-run initiation. Customer-facing services remain on legacy during parallel period. Deliverable: Core banking operational in cloud with parallel validation.",
        "Quarter 5-6 (January-June 2028): Legacy decommissioning, advanced analytics deployment, mobile application launch, fraud detection system activation. Deliverable: Full cloud-native operation with all advanced capabilities.",
    ],
    "Budget and Resource Allocation": [
        "The total program budget of $4.7 million is distributed across infrastructure costs ($1.8M), professional services ($1.4M), software licensing ($800K), training and change management ($400K), and contingency ($300K). Monthly cloud operating costs are estimated at $85,000, compared to the current $142,000 in data center and maintenance expenses.",
        "Personnel requirements include a dedicated migration team of 8 full-time equivalents: 2 cloud architects, 3 senior developers, 1 database specialist, 1 security engineer, and 1 project manager. An additional 4 consultants from our selected systems integrator will supplement the team during peak migration periods in Phases 1 and 2.",
        "Return on investment projections show break-even at month 34 post-initiation. Annual savings of $1.2 million from reduced infrastructure costs, lower maintenance overhead, and improved operational efficiency compound to a projected 5-year net benefit of $3.8 million.",
        "Capital expenditure requirements are front-loaded, with 60% of total spending occurring in the first 12 months. We recommend a phased procurement approach to spread budget impact across fiscal years, with the option to accelerate spending if early phases demonstrate positive outcomes.",
    ],
    "Expected Outcomes and KPIs": [
        "System performance improvements are projected to reduce average transaction processing time from 3.2 seconds to under 400 milliseconds. The new architecture supports horizontal scaling, enabling the platform to handle 10x current peak transaction volumes without degradation.",
        "Security posture will be significantly enhanced through automated vulnerability scanning, real-time threat detection, and zero-trust network architecture. The target is zero critical vulnerabilities in the production environment, with automated patching reducing the current 14-day patch cycle to under 24 hours for critical updates.",
        "Customer satisfaction metrics, as measured by the existing Net Promoter Score survey, are expected to improve from the current 32 to a target of 55 within 12 months of the mobile banking launch. Digital channel adoption is projected to increase from 41% to 75% of active customers.",
        "Operational efficiency gains include a 40% reduction in incident response time, 99.99% system availability (up from current 99.7%), and a 60% reduction in manual processing tasks through automation of routine operations.",
    ],
    "Appendix: Technical Specifications": [
        "Cloud Architecture: Multi-region deployment across AWS us-east-1 (primary) and us-west-2 (disaster recovery). VPC configuration with public, private, and isolated subnets. Transit Gateway for inter-VPC communication. AWS Shield Advanced and WAF for DDoS protection and web application security.",
        "Container Platform: Amazon EKS with Kubernetes 1.28, Istio service mesh for inter-service communication, Prometheus and Grafana for observability. CI/CD pipeline using GitHub Actions with ArgoCD for GitOps-based deployments. Container images stored in Amazon ECR with automated vulnerability scanning.",
        "Database Layer: Amazon Aurora PostgreSQL 15.4 with read replicas across availability zones. DynamoDB for session management and caching. Amazon ElastiCache (Redis) for application-level caching. Data migration via AWS Database Migration Service with ongoing replication during parallel-run.",
        "Security Stack: AWS IAM with SAML 2.0 federation to existing Active Directory. Amazon Cognito for customer authentication with MFA support. AWS CloudTrail and GuardDuty for audit logging and threat detection. HashiCorp Vault for secrets management. Encryption at rest (AES-256) and in transit (TLS 1.3).",
        "Monitoring and Operations: Amazon CloudWatch for infrastructure monitoring. PagerDuty integration for incident management. Datadog APM for application performance monitoring. ELK stack (Elasticsearch, Logstash, Kibana) for centralized log aggregation and analysis. Automated runbooks via AWS Systems Manager.",
    ],
}

# ── Revised content: text changes ───────────────────────────────────────
# 3 paragraphs rewritten + 2 sentences added

SECTION_CONTENT_REVISED = {}
for k, v in SECTION_CONTENT_ORIGINAL.items():
    SECTION_CONTENT_REVISED[k] = list(v)  # shallow copy

# Rewrite 1: Executive Summary paragraph 1
SECTION_CONTENT_REVISED["Executive Summary"][0] = (
    "Meridian Financial Group's current technology infrastructure, originally deployed in 2014, "
    "has reached end-of-life across multiple critical components. Rising maintenance costs, "
    "an expanding attack surface, and inability to support modern API-driven integrations "
    "make a comprehensive modernization initiative not merely desirable but urgently necessary "
    "for the organization's continued competitive viability in the evolving financial services landscape."
)

# Rewrite 2: Risk Analysis paragraph 2
SECTION_CONTENT_REVISED["Risk Analysis and Mitigation"][1] = (
    "Maintaining seamless operations during the parallel-run window demands meticulous capacity "
    "engineering. We have provisioned 40% additional compute headroom — up from the initially "
    "estimated 30% — after stress-testing revealed higher-than-expected memory consumption during "
    "dual-system reconciliation. Branch connectivity will be safeguarded through SD-WAN failover "
    "paths with automatic traffic rerouting in under 500 milliseconds."
)

# Rewrite 3: Budget paragraph 3
SECTION_CONTENT_REVISED["Budget and Resource Allocation"][2] = (
    "Updated financial modeling shows break-even at month 30 rather than the originally projected "
    "month 34, primarily due to accelerated decommissioning of legacy hardware and earlier-than-expected "
    "reductions in third-party maintenance contracts. The revised 5-year net benefit stands at $4.3 million, "
    "an improvement of $500K over initial estimates."
)

# Added sentence 1: at the end of "Proposed Migration Strategy" paragraph 4
SECTION_CONTENT_REVISED["Proposed Migration Strategy"][3] += (
    " Additionally, a dedicated accessibility review will be conducted during this phase to ensure "
    "compliance with WCAG 2.2 AA standards across all customer-facing digital touchpoints."
)

# Added sentence 2: at the end of "Expected Outcomes and KPIs" paragraph 3
SECTION_CONTENT_REVISED["Expected Outcomes and KPIs"][2] += (
    " Furthermore, we anticipate a 25% reduction in call center volume as self-service capabilities "
    "expand through the new mobile and web platforms."
)


def build_document(doc, section_content, heading_font_size, body_space_after):
    """Build the full proposal document with given content and formatting."""
    # Title page
    title_para = doc.add_heading(TITLE, level=0)
    for run in title_para.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub_para = doc.add_paragraph(SUBTITLE)
    sub_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in sub_para.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    date_para = doc.add_paragraph(DATE_LINE)
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in date_para.runs:
        run.font.size = Pt(12)
        run.font.italic = True

    # Confidentiality notice
    conf = doc.add_paragraph()
    conf.paragraph_format.space_before = Pt(36)
    conf.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = conf.add_run("CONFIDENTIAL — For Internal Use Only")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
    run.bold = True

    doc.add_page_break()

    # Table of Contents placeholder
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for run in toc_heading.runs:
        run.font.size = heading_font_size
    for i, section_name in enumerate(SECTIONS, 1):
        toc_entry = doc.add_paragraph(f"{i}. {section_name}")
        toc_entry.paragraph_format.space_after = Pt(4)
        for run in toc_entry.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # Sections
    for i, section_name in enumerate(SECTIONS):
        # Heading with formatting
        h = doc.add_heading(section_name, level=1)
        for run in h.runs:
            run.font.size = heading_font_size

        # Body paragraphs
        for para_text in section_content[section_name]:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = body_space_after
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = "Liberation Serif"

        # Add page break after most sections (to reach ~8 pages)
        if i < len(SECTIONS) - 1:
            doc.add_page_break()


def create_initial():
    os.makedirs(f'{WORKDIR}/Documents', exist_ok=True)

    # ── Create the original document ──
    doc_orig = Document()

    # Page setup
    section = doc_orig.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    build_document(doc_orig, SECTION_CONTENT_ORIGINAL,
                   HEADING_FONT_ORIGINAL, BODY_SPACE_AFTER_ORIGINAL)

    doc_orig.save(OUTPUT)
    print(f'Original document created: {OUTPUT}')

    # ── Create the revised document ──
    doc_rev = Document()

    section_rev = doc_rev.sections[0]
    section_rev.page_width = Inches(8.5)
    section_rev.page_height = Inches(11)
    section_rev.left_margin = Inches(1.25)
    section_rev.right_margin = Inches(1.25)
    section_rev.top_margin = Inches(1.0)
    section_rev.bottom_margin = Inches(1.0)

    # Revised has: different heading font sizes (4 headings changed)
    # and different paragraph spacing (6 paragraphs changed)
    # But same text changes applied
    build_document(doc_rev, SECTION_CONTENT_REVISED,
                   HEADING_FONT_REVISED, BODY_SPACE_AFTER_REVISED)

    doc_rev.save(REVISED)
    print(f'Revised document created: {REVISED}')

    # ── Open original in LibreOffice Writer ──
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
