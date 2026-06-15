"""
Initial Setup: poem_submission.docx — pre-task state for writer_creative_060
Task ID: writer_creative_060
Domain: libreoffice_writer

Creates a poem document with:
- Title 'The Weight of Winter' (left-aligned, 12pt Times New Roman, unformatted)
- Byline 'by Amara Johnson' (left-aligned, 12pt Times New Roman)
- 3 stanzas of 4 lines each (12 lines total), single-spaced, left-aligned
- Stanzas separated by single blank lines
- NO contact info block
- NO bold title, NO Liberation Sans, NO centering, NO double-spacing
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'poem_submission'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default styles / margins — keep default page setup
    # Set default paragraph style to Times New Roman 12pt for this doc
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- Title: "The Weight of Winter" ---
    # Left-aligned, 12pt Times New Roman, no bold, no special font — pre-task state
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_para.paragraph_format.line_spacing = 1.0
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    title_run = title_para.add_run('The Weight of Winter')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_run.bold = False

    # --- Byline: "by Amara Johnson" ---
    byline_para = doc.add_paragraph()
    byline_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    byline_para.paragraph_format.line_spacing = 1.0
    byline_para.paragraph_format.space_before = Pt(0)
    byline_para.paragraph_format.space_after = Pt(0)
    byline_run = byline_para.add_run('by Amara Johnson')
    byline_run.font.name = 'Times New Roman'
    byline_run.font.size = Pt(12)

    # --- Blank line separator before poem ---
    blank = doc.add_paragraph()
    blank.paragraph_format.line_spacing = 1.0
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after = Pt(0)

    # --- Poem stanzas (3 stanzas, 4 lines each) ---
    stanzas = [
        [
            'The frost comes early to the window pane,',
            'It lays a silence on the sleeping ground,',
            'The trees stand bare, their branches full of rain,',
            'And winter wraps the world without a sound.',
        ],
        [
            'I watch the gray clouds settle on the hill,',
            'The river slows beneath its coat of ice,',
            'The fields lie fallow, hushed and cold and still,',
            'As if the earth has paid its season\'s price.',
        ],
        [
            'Yet even now, beneath the frozen earth,',
            'The seeds are waiting for the warming sun,',
            'They hold inside them promises of birth,',
            'And know that winter\'s work is never done.',
        ],
    ]

    for stanza_idx, stanza in enumerate(stanzas):
        # Add each line of the stanza
        for line in stanza:
            line_para = doc.add_paragraph()
            line_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            line_para.paragraph_format.line_spacing = 1.0
            line_para.paragraph_format.space_before = Pt(0)
            line_para.paragraph_format.space_after = Pt(0)
            line_run = line_para.add_run(line)
            line_run.font.name = 'Times New Roman'
            line_run.font.size = Pt(12)

        # Add single blank line after each stanza (except last)
        if stanza_idx < len(stanzas) - 1:
            blank_para = doc.add_paragraph()
            blank_para.paragraph_format.line_spacing = 1.0
            blank_para.paragraph_format.space_before = Pt(0)
            blank_para.paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
