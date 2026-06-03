"""
Initial Setup: Business proposal document with single line spacing throughout
Task ID: wrpara_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER

WORKDIR = '/home/user'
TASK_ID = 'wrpara_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ========== PAGE 1: COVER PAGE ==========
    # Add some vertical spacing before title
    for _ in range(6):
        spacer = doc.add_paragraph('')
        spacer.paragraph_format.line_spacing = 1.0
        spacer.paragraph_format.space_after = Pt(0)

    # Title
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.line_spacing = 1.0
    title_para.paragraph_format.space_after = Pt(0)
    title_run = title_para.add_run('Strategic Partnership Proposal')
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.line_spacing = 1.0
    subtitle_para.paragraph_format.space_after = Pt(0)
    subtitle_run = subtitle_para.add_run('Prepared for Acme Corp')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    subtitle_run.italic = True

    # Date line on cover
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.line_spacing = 1.0
    date_para.paragraph_format.space_before = Pt(24)
    date_run = date_para.add_run('March 2026')
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break after cover
    doc.add_page_break()

    # ========== PAGE 2: TABLE OF CONTENTS ==========
    toc_heading = doc.add_paragraph()
    toc_heading.paragraph_format.line_spacing = 1.0
    toc_heading.paragraph_format.space_after = Pt(12)
    toc_run = toc_heading.add_run('Table of Contents')
    toc_run.bold = True
    toc_run.font.size = Pt(16)
    toc_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    toc_entries = [
        ('Executive Summary', '3'),
        ('1. Company Overview', '4'),
        ('2. Market Analysis', '5'),
        ('3. Partnership Objectives', '6'),
        ('4. Implementation Plan', '7'),
        ('5. Financial Projections', '8'),
        ('Appendix A: Supporting Data', '9'),
        ('Appendix B: Team Biographies', '10'),
    ]

    for entry_text, page_num in toc_entries:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.line_spacing = 1.0
        toc_para.paragraph_format.space_after = Pt(2)
        # Add tab stop with dot leader at right margin
        tab_stops = toc_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        toc_para.add_run(entry_text)
        toc_para.add_run(f'\t{page_num}')

    # Page break after TOC
    doc.add_page_break()

    # ========== PAGE 3: EXECUTIVE SUMMARY ==========
    exec_heading = doc.add_paragraph()
    exec_heading.paragraph_format.line_spacing = 1.0
    exec_heading.paragraph_format.space_after = Pt(8)
    exec_run = exec_heading.add_run('Executive Summary')
    exec_run.bold = True
    exec_run.font.size = Pt(16)
    exec_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    exec_para1 = doc.add_paragraph(
        'This proposal outlines a strategic partnership between Meridian Dynamics and Acme Corp '
        'to jointly develop next-generation supply chain optimization solutions. With combined '
        'expertise in artificial intelligence and logistics infrastructure, this partnership '
        'positions both organizations to capture an estimated $4.2 billion market opportunity '
        'in the enterprise supply chain automation sector by 2028. Our preliminary analysis '
        'indicates a potential 35% reduction in operational costs for target customers, with '
        'projected first-year revenue of $12.8 million from the joint venture.'
    )
    exec_para1.paragraph_format.line_spacing = 1.0
    exec_para1.paragraph_format.space_after = Pt(6)

    exec_para2 = doc.add_paragraph(
        'The partnership leverages Meridian Dynamics\u2019 proprietary machine learning algorithms '
        'and Acme Corp\u2019s extensive distribution network spanning 47 countries. Initial pilot '
        'programs with three Fortune 500 clients have demonstrated a 28% improvement in delivery '
        'times and a 19% reduction in inventory carrying costs. This proposal details the '
        'governance structure, investment requirements of $6.5 million over 18 months, and a '
        'phased implementation timeline with clearly defined milestones and success metrics.'
    )
    exec_para2.paragraph_format.line_spacing = 1.0
    exec_para2.paragraph_format.space_after = Pt(6)

    # Page break after executive summary
    doc.add_page_break()

    # ========== PAGES 4-8: MAIN BODY (5 sections, 3 paragraphs each) ==========

    sections_data = [
        {
            'title': '1. Company Overview',
            'paragraphs': [
                'Meridian Dynamics was founded in 2018 by Dr. Elena Vasquez and James Thornton with a mission to transform enterprise logistics through cutting-edge artificial intelligence. Headquartered in San Francisco with offices in London, Singapore, and Sao Paulo, the company has grown to 340 employees and serves over 120 enterprise clients globally. Our flagship product, MeridianFlow, processes over 2.3 million supply chain transactions daily.',
                'Acme Corp, established in 1987, has built one of the most extensive distribution networks in the industry, covering 47 countries with 215 distribution centers. Under the leadership of CEO Patricia Hawkins, Acme Corp generated $3.8 billion in revenue last fiscal year, with a compound annual growth rate of 12% over the past five years. The company\u2019s infrastructure handles approximately 890,000 shipments per day.',
                'Together, both organizations bring complementary strengths that create a unique competitive advantage. Meridian\u2019s AI-driven optimization technology combined with Acme\u2019s physical infrastructure and established client relationships forms the foundation of a partnership that neither competitor can easily replicate in the current market landscape.',
            ]
        },
        {
            'title': '2. Market Analysis',
            'paragraphs': [
                'The global supply chain management market is projected to reach $37.4 billion by 2027, growing at a CAGR of 11.2% according to Grand View Research. Key growth drivers include increasing complexity of global trade networks, rising consumer expectations for faster delivery, and the accelerating adoption of Industry 4.0 technologies across manufacturing and retail sectors.',
                'Our competitive analysis identifies three primary market segments: enterprise logistics optimization valued at $14.6 billion, warehouse automation at $9.2 billion, and last-mile delivery solutions at $8.1 billion. The proposed partnership targets the enterprise logistics segment, where AI-driven solutions currently represent only 18% of deployments, indicating significant growth potential.',
                'Regional analysis reveals that North America and Europe account for 62% of current spending, but the Asia-Pacific region is experiencing the fastest growth at 15.8% CAGR. Our combined presence in these regions positions the partnership to capture share across all major geographies, with particular strength in emerging markets where Acme Corp has established relationships with local distributors.',
            ]
        },
        {
            'title': '3. Partnership Objectives',
            'paragraphs': [
                'The primary objective of this strategic partnership is to develop and deploy an integrated supply chain intelligence platform that combines Meridian\u2019s predictive analytics with Acme\u2019s real-time logistics data. The platform, tentatively named "Nexus," will offer end-to-end visibility across procurement, warehousing, transportation, and delivery operations for enterprise clients.',
                'Secondary objectives include establishing a joint research laboratory in Austin, Texas, with an initial team of 25 data scientists and engineers focused on advancing multi-modal optimization algorithms. The lab will also collaborate with Stanford University and MIT on foundational research in reinforcement learning applications for dynamic route optimization.',
                'Revenue targets for the partnership include achieving $12.8 million in first-year revenue from the joint platform, scaling to $45 million by year three, and reaching profitability by the end of the second fiscal year. These targets are supported by a pipeline of 34 qualified enterprise prospects and letters of intent from seven existing clients representing combined annual logistics spend of $890 million.',
            ]
        },
        {
            'title': '4. Implementation Plan',
            'paragraphs': [
                'Phase 1 (Months 1-6) focuses on technical integration and pilot deployment. Engineering teams from both organizations will work to integrate Meridian\u2019s API endpoints with Acme\u2019s existing warehouse management and transportation management systems. Three pilot customers \u2013 Novatel Industries, Brightpath Retail, and Consolidated Freight \u2013 have agreed to participate in the beta program.',
                'Phase 2 (Months 7-12) involves scaling the platform to support 50 concurrent enterprise clients and expanding the feature set to include predictive demand forecasting, automated procurement recommendations, and real-time carbon footprint tracking. A dedicated customer success team of 12 specialists will be recruited and trained during this phase.',
                'Phase 3 (Months 13-18) targets full commercial launch with a comprehensive go-to-market strategy spanning direct sales, channel partnerships, and a self-service tier for mid-market clients. Marketing investment of $2.1 million will support brand awareness campaigns, industry conference sponsorships, and a thought leadership content program.',
            ]
        },
        {
            'title': '5. Financial Projections',
            'paragraphs': [
                'The total investment required for the partnership is $6.5 million over 18 months, split equally between both organizations. This includes $2.8 million for technology development, $1.9 million for personnel and recruitment, $1.2 million for infrastructure and cloud computing costs, and $600,000 for legal, compliance, and administrative expenses.',
                'Revenue projections are based on a tiered SaaS pricing model with annual contract values ranging from $180,000 for standard deployments to $750,000 for enterprise-plus configurations. Gross margins are expected to reach 72% by year two as the platform matures and operational efficiencies are realized through automation of customer onboarding processes.',
                'A sensitivity analysis across three scenarios (conservative, base, and optimistic) yields NPV estimates ranging from $18.2 million to $47.6 million over a five-year horizon, with the base case projecting an IRR of 34% and payback period of 26 months. Risk factors include potential delays in enterprise sales cycles and competitive responses from established supply chain vendors.',
            ]
        },
    ]

    for i, section in enumerate(sections_data):
        # Section heading
        sec_heading = doc.add_paragraph()
        sec_heading.paragraph_format.line_spacing = 1.0
        sec_heading.paragraph_format.space_before = Pt(12) if i > 0 else Pt(0)
        sec_heading.paragraph_format.space_after = Pt(8)
        sec_run = sec_heading.add_run(section['title'])
        sec_run.bold = True
        sec_run.font.size = Pt(14)
        sec_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        for para_text in section['paragraphs']:
            body_para = doc.add_paragraph(para_text)
            body_para.paragraph_format.line_spacing = 1.0
            body_para.paragraph_format.space_after = Pt(6)

        # Page break after each section except the last
        if i < len(sections_data) - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
