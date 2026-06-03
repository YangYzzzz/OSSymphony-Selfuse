"""
Initial Setup: Employee Exit Checklist - bare document with title and employee info fields only.
Task ID: writer_hr_056
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading('Employee Exit Checklist', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle / company info ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # --- Blank separator ---
    doc.add_paragraph()

    # --- Employee info fields ---
    p1 = doc.add_paragraph()
    run1 = p1.add_run('Employee Name: ')
    run1.bold = True
    run1.font.size = Pt(12)
    run1b = p1.add_run('_______________________________')
    run1b.font.size = Pt(12)

    p2 = doc.add_paragraph()
    run2 = p2.add_run('Department: ')
    run2.bold = True
    run2.font.size = Pt(12)
    run2b = p2.add_run('_______________________________')
    run2b.font.size = Pt(12)

    p3 = doc.add_paragraph()
    run3 = p3.add_run('Last Day of Employment: ')
    run3.bold = True
    run3.font.size = Pt(12)
    run3b = p3.add_run('___________________')
    run3b.font.size = Pt(12)

    p4 = doc.add_paragraph()
    run4 = p4.add_run('Manager: ')
    run4.bold = True
    run4.font.size = Pt(12)
    run4b = p4.add_run('_______________________________')
    run4b.font.size = Pt(12)

    # --- Instructions paragraph ---
    doc.add_paragraph()
    instructions = doc.add_paragraph()
    run_inst = instructions.add_run(
        'This checklist must be completed by the respective departments prior to the '
        "employee's last day. Each responsible party should initial and date upon completion "
        'of their assigned tasks.'
    )
    run_inst.font.size = Pt(11)
    run_inst.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
