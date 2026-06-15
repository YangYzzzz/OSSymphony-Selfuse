"""
Reward Script: Employee Exit Checklist Document
Task ID: writer_hr_056
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Four department section headings present
  Component 2 (0.30): Four tables with correct 4-column structure (Task, Completed, Date, Initials)
  Component 3 (0.25): Each table has 3-5 task rows (at least 13 total data rows)
  Component 4 (0.20): Checkbox characters in the Completed column of data rows
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_056'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Four department section headings (0.25 points)
    # The golden doc should have headings for IT Department, HR Department, Facilities, Manager
    # Only count paragraphs that are headings or whose text IS the section name (not field labels like "Manager: ___")
    try:
        required_sections = ['it department', 'hr department', 'facilities', 'manager']
        found_sections = []
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            is_heading = para.style.name.lower().startswith('heading')
            for section in required_sections:
                # Must be a heading style OR the paragraph text is essentially just the section name
                # Exclude lines like "Manager: ___" (contain colon or underscore after the keyword)
                if section in text_lower and section not in found_sections:
                    if is_heading or text_lower == section:
                        found_sections.append(section)

        sections_found_count = len(found_sections)
        if sections_found_count == 4:
            print(f"PASS: Component 1 — All 4 section headings found: {found_sections} (0.25 pts)")
            total_score += 0.25
        elif sections_found_count > 0:
            partial = 0.25 * (sections_found_count / 4)
            print(f"PARTIAL: Component 1 — {sections_found_count}/4 sections found: {found_sections} ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No department section headings found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Four tables with 4-column structure (Task, Completed, Date, Initials) (0.30 points)
    # Each table header row must have these four columns
    try:
        tables = doc.tables
        num_tables = len(tables)
        valid_tables = 0
        expected_headers = ['task', 'completed', 'date', 'initials']

        for t_idx, table in enumerate(tables):
            if len(table.columns) >= 4 and len(table.rows) >= 2:
                # Check header row
                header_cells = [table.cell(0, c).text.strip().lower() for c in range(4)]
                if all(eh in header_cells for eh in expected_headers):
                    valid_tables += 1
                    print(f"  Table {t_idx}: valid structure, headers={header_cells}")
                else:
                    print(f"  Table {t_idx}: invalid headers={header_cells}")
            else:
                print(f"  Table {t_idx}: insufficient cols={len(table.columns)} or rows={len(table.rows)}")

        if valid_tables >= 4:
            print(f"PASS: Component 2 — {valid_tables} valid tables found (0.30 pts)")
            total_score += 0.30
        elif valid_tables > 0:
            partial = 0.30 * (valid_tables / 4)
            print(f"PARTIAL: Component 2 — {valid_tables}/4 valid tables ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No valid tables found (total tables: {num_tables})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Each table has 3-5 task rows, total >= 13 data rows (0.25 points)
    # Data rows = all rows except header row
    try:
        tables = doc.tables
        total_data_rows = 0
        tables_with_valid_row_count = 0

        for t_idx, table in enumerate(tables):
            data_rows = len(table.rows) - 1  # exclude header
            if 3 <= data_rows <= 5:
                tables_with_valid_row_count += 1
            total_data_rows += data_rows
            print(f"  Table {t_idx}: {data_rows} data rows")

        if num_tables >= 4 and total_data_rows >= 13 and tables_with_valid_row_count >= 4:
            print(f"PASS: Component 3 — {total_data_rows} total data rows across {num_tables} tables, all with 3-5 rows (0.25 pts)")
            total_score += 0.25
        elif num_tables >= 4 and total_data_rows >= 10:
            partial = 0.25 * 0.6
            print(f"PARTIAL: Component 3 — {total_data_rows} data rows, {tables_with_valid_row_count}/4 tables valid ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — tables={num_tables}, total data rows={total_data_rows}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Checkbox characters in Completed column of data rows (0.20 points)
    # The Completed column (col index 1) should have checkbox characters (☐, ☑, □, ✓, ✗, etc.)
    try:
        tables = doc.tables
        checkbox_chars = {'☐', '☑', '□', '■', '✓', '✗', '✔', '✘', '☒', '▢', '▣'}
        rows_with_checkbox = 0
        total_check_rows = 0

        for t_idx, table in enumerate(tables):
            for r_idx in range(1, len(table.rows)):  # skip header
                total_check_rows += 1
                cell_text = table.cell(r_idx, 1).text.strip()
                if any(c in cell_text for c in checkbox_chars):
                    rows_with_checkbox += 1

        if total_check_rows > 0 and rows_with_checkbox >= total_check_rows * 0.8:
            print(f"PASS: Component 4 — {rows_with_checkbox}/{total_check_rows} data rows have checkbox characters (0.20 pts)")
            total_score += 0.20
        elif rows_with_checkbox > 0:
            ratio = rows_with_checkbox / max(total_check_rows, 1)
            partial = 0.20 * ratio
            print(f"PARTIAL: Component 4 — {rows_with_checkbox}/{total_check_rows} have checkboxes ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No checkbox characters found in Completed column")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
