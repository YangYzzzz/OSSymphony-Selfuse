"""
Initial Setup: Create a business review document with empty document properties.
Task ID: writer_biz_055
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document properties: intentionally left empty/default ---
    # The task requires the user to set Title, Subject, and Author.
    # We ensure these are blank/default in the initial state.
    doc.core_properties.title = ''
    doc.core_properties.subject = ''
    doc.core_properties.author = ''

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title heading ---
    heading = doc.add_heading('Annual Business Review', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report provides a comprehensive overview of Meridian Technologies\' '
        'performance during the fiscal year 2024-2025. The company has demonstrated '
        'strong growth across its core business segments, with particular success in '
        'cloud infrastructure services and enterprise software solutions.'
    )
    doc.add_paragraph(
        'Key highlights include a 23% year-over-year revenue increase, successful '
        'expansion into three new international markets, and the launch of our '
        'next-generation analytics platform. Despite headwinds in global supply chains, '
        'the organization maintained healthy profit margins and exceeded quarterly targets '
        'in three of four reporting periods.'
    )

    # --- Financial Overview ---
    doc.add_heading('Financial Overview', level=1)
    doc.add_paragraph(
        'Total revenue for fiscal year 2024-2025 reached $287.4 million, up from '
        '$233.7 million the prior year. Operating income improved to $48.2 million, '
        'reflecting a 16.8% operating margin. The company\'s balance sheet remains robust '
        'with $92.1 million in cash and short-term investments as of March 31, 2025.'
    )

    # Financial summary table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Metric', 'FY 2024', 'FY 2025']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    data = [
        ['Total Revenue', '$233.7M', '$287.4M'],
        ['Operating Income', '$38.6M', '$48.2M'],
        ['Net Profit', '$29.1M', '$37.8M'],
        ['Gross Margin', '62.4%', '64.1%'],
        ['Headcount', '1,847', '2,234'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Market Analysis ---
    doc.add_heading('Market Analysis', level=1)
    doc.add_paragraph(
        'The global enterprise software market continued its rapid expansion in 2025, '
        'with cloud adoption accelerating across industries. Meridian Technologies '
        'strengthened its competitive position through strategic partnerships with Amazon '
        'Web Services and Microsoft Azure, enabling integrated deployment options for '
        'enterprise clients.'
    )
    doc.add_paragraph(
        'Notable client acquisitions during the fiscal year include Deutsche Telekom, '
        'Sumitomo Corporation, and Brookfield Asset Management. These partnerships '
        'represent combined annual recurring revenue of approximately $14.8 million.'
    )

    # --- Strategic Initiatives ---
    doc.add_heading('Strategic Initiatives', level=1)
    doc.add_paragraph(
        'The company pursued several strategic initiatives during the review period:'
    )
    items = [
        'Launch of MeridianAI Analytics Platform with predictive modeling capabilities',
        'Expansion of Singapore and Frankfurt data center facilities',
        'Acquisition of CloudBridge Solutions for $42.5 million',
        'Implementation of enhanced cybersecurity framework across all products',
        'Establishment of the Meridian Innovation Lab in Austin, Texas',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # --- Human Resources ---
    doc.add_heading('Human Resources & Organizational Development', level=1)
    doc.add_paragraph(
        'Meridian Technologies invested heavily in talent acquisition and retention '
        'throughout fiscal year 2025. The engineering team grew by 28%, with key hires '
        'including Dr. Priya Sharma as Chief Technology Officer and James Rodriguez as '
        'VP of Product Engineering. Employee satisfaction scores improved to 4.3 out of '
        '5.0, driven by expanded remote work policies and the introduction of a '
        'comprehensive professional development program.'
    )

    # --- Outlook ---
    doc.add_heading('Outlook for FY 2026', level=1)
    doc.add_paragraph(
        'Looking ahead, Meridian Technologies is well-positioned for continued growth. '
        'The company targets revenue of $340-360 million for fiscal year 2026, '
        'representing 18-25% growth. Key priorities include deepening AI/ML integration '
        'across the product portfolio, expanding presence in the Asia-Pacific region, and '
        'achieving SOC 2 Type II certification for all cloud services by Q3 2026.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
