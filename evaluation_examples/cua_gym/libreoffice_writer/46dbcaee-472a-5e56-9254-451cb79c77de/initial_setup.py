"""
Initial Setup: Turn off track changes recording in the proposal document.
Task ID: writer_rm_002
Domain: libreoffice_writer

Creates a Q4 Proposal document with track changes recording enabled
and 5 tracked insertions visible in the document.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
from datetime import datetime

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

AUTHOR = 'Sarah Chen'
REVISION_DATE = '2025-11-15T10:30:00Z'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Q4 2025 Strategic Growth Proposal', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by the Business Development Team')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run.italic = True

    doc.add_paragraph()  # spacer

    # --- Section 1: Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    p1 = doc.add_paragraph(
        'This proposal outlines our strategic plan for the fourth quarter of 2025. '
        'Our primary objectives focus on expanding market presence in the Asia-Pacific '
        'region, launching two new product lines, and strengthening our partnership '
        'network with key industry players.'
    )

    p2 = doc.add_paragraph(
        'The projected investment of $2.4 million is expected to yield a 34% increase '
        'in regional revenue by Q1 2026. This aligns with the board-approved growth '
        'targets established during the annual planning session in March.'
    )

    # --- Section 2: Market Analysis ---
    doc.add_heading('2. Market Analysis', level=1)
    p3 = doc.add_paragraph(
        'Current market research indicates a significant shift toward cloud-based '
        'solutions across enterprise customers in our target demographics. Competitors '
        'such as Meridian Technologies and Apex Solutions have already begun pivoting '
        'their offerings, increasing urgency for our own transition.'
    )

    p4 = doc.add_paragraph(
        'Key market indicators from the Q3 industry report show:'
    )

    # Bullet points
    bullets = [
        'Enterprise cloud adoption rate: 67% (up from 52% in Q2)',
        'Average contract value in APAC: $185,000 per annum',
        'Customer retention rate for cloud-first vendors: 91%',
        'Projected SaaS market growth: 22% year-over-year',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # --- Section 3: Proposed Initiatives ---
    doc.add_heading('3. Proposed Initiatives', level=1)

    doc.add_heading('3.1 Product Launch: CloudSync Pro', level=2)
    p5 = doc.add_paragraph(
        'CloudSync Pro represents our flagship cloud migration tool, designed to '
        'simplify the transition from on-premise infrastructure. Development is on '
        'track with a target launch date of November 15, 2025.'
    )

    doc.add_heading('3.2 Regional Expansion: Singapore Office', level=2)
    p6 = doc.add_paragraph(
        'We recommend establishing a satellite office in Singapore to serve as the '
        'APAC operations hub. Initial staffing requirements include a regional director, '
        'three senior account managers, and two technical consultants.'
    )

    doc.add_heading('3.3 Partnership Program', level=2)
    p7 = doc.add_paragraph(
        'The proposed Platinum Partner Program will offer tiered incentives to '
        'resellers and system integrators. Early discussions with DataFlow Systems '
        'and NorthBridge Consulting indicate strong interest in formalized partnerships.'
    )

    # --- Section 4: Budget ---
    doc.add_heading('4. Budget Overview', level=1)

    # Budget table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Initiative', 'Estimated Cost', 'Timeline']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    budget_data = [
        ['CloudSync Pro Development', '$850,000', 'Oct - Dec 2025'],
        ['Singapore Office Setup', '$620,000', 'Nov 2025 - Jan 2026'],
        ['Marketing Campaign (APAC)', '$380,000', 'Oct - Dec 2025'],
        ['Partnership Program Launch', '$290,000', 'Nov - Dec 2025'],
        ['Contingency Reserve', '$260,000', 'As needed'],
    ]
    for r, row_data in enumerate(budget_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    # --- Section 5: Timeline ---
    doc.add_heading('5. Implementation Timeline', level=1)
    p8 = doc.add_paragraph(
        'Phase 1 (October 2025): Finalize CloudSync Pro beta testing, begin '
        'Singapore office lease negotiations, and launch partner recruitment campaign.'
    )
    p9 = doc.add_paragraph(
        'Phase 2 (November 2025): Product launch event, office setup completion, '
        'and first wave of partner onboarding with training sessions.'
    )
    p10 = doc.add_paragraph(
        'Phase 3 (December 2025): Full operational capacity, performance review, '
        'and Q1 2026 planning based on initial results and market feedback.'
    )

    # --- Section 6: Conclusion ---
    doc.add_heading('6. Conclusion', level=1)
    p_conc = doc.add_paragraph(
        'This proposal represents a significant but measured investment in our '
        'company\'s future growth. The combination of product innovation, geographic '
        'expansion, and strategic partnerships positions us to capture a meaningful '
        'share of the rapidly growing APAC cloud services market.'
    )

    # ---- Now add 5 tracked insertions via XML manipulation ----
    # Track changes in OOXML use <w:ins> elements wrapping <w:r> runs
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def make_ins_element(text, rev_id, author=AUTHOR, date=REVISION_DATE):
        """Create a w:ins element containing a run with the given text."""
        ins = etree.SubElement(
            etree.Element('dummy'),  # temporary parent
            qn('w:ins'),
            {
                qn('w:id'): str(rev_id),
                qn('w:author'): author,
                qn('w:date'): date,
            }
        )
        r = etree.SubElement(ins, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        t = etree.SubElement(r, qn('w:t'))
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        return ins

    # Insertion 1: Add "and scalability" to executive summary paragraph (p1)
    ins1 = make_ins_element(' and scalability', rev_id=1)
    # Append to last run's parent paragraph
    p1_elem = p1._element
    p1_elem.append(ins1)

    # Insertion 2: Add emphasis text to market analysis (p3)
    ins2 = make_ins_element(' — a critical competitive advantage', rev_id=2)
    p3_elem = p3._element
    p3_elem.append(ins2)

    # Insertion 3: Add note about staffing to Singapore section (p6)
    ins3 = make_ins_element(' Additional hires may be required in Q1 2026 based on demand.', rev_id=3)
    p6_elem = p6._element
    p6_elem.append(ins3)

    # Insertion 4: Add risk mention to Phase 3 timeline (p10)
    ins4 = make_ins_element(' Risk mitigation strategies will be reviewed at this stage.', rev_id=4)
    p10_elem = p10._element
    p10_elem.append(ins4)

    # Insertion 5: Add stakeholder note to conclusion
    ins5 = make_ins_element(' We recommend presenting this to the full board at the October meeting.', rev_id=5)
    p_conc_elem = p_conc._element
    p_conc_elem.append(ins5)

    # ---- Enable track changes recording in document settings ----
    # Access the settings part and add w:trackChanges element
    settings_part = doc.settings.element
    # Add trackChanges element to indicate recording is ON
    # In OOXML, <w:trackChanges/> in settings.xml means recording is enabled
    track_changes_elem = etree.SubElement(settings_part, qn('w:trackChanges'))

    # Also set rsid for revision tracking
    rsids = settings_part.find(qn('w:rsids'))
    if rsids is None:
        rsids = etree.SubElement(settings_part, qn('w:rsids'))
    rsid_root = etree.SubElement(rsids, qn('w:rsidRoot'), {qn('w:val'): '00A12B3C'})

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')

create_initial()
