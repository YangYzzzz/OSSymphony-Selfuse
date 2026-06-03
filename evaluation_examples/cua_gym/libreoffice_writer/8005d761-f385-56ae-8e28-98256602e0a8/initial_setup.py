"""
Initial Setup: Technical Manual with 'API endpoint' in plain text
Task ID: osworld_writer_text_formatting_basic_004
Domain: libreoffice_writer

Creates a 5-page technical manual document where 'API endpoint' appears
8 times across 3 sections in plain black text (no bold/underline/color).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_text_formatting_basic_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_plain_paragraph(doc, text, style=None, bold_heading=False, size=None):
    """Add a paragraph with plain black text (no special formatting)."""
    if style:
        para = doc.add_paragraph(style=style)
    else:
        para = doc.add_paragraph()
    run = para.add_run(text)
    if bold_heading:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    # Ensure plain black color
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return para


def add_mixed_paragraph(doc, parts):
    """
    Add a paragraph with mixed formatting segments.
    parts is a list of (text, is_keyword) tuples.
    Keywords ('API endpoint') are plain black - no bold/underline/color.
    """
    para = doc.add_paragraph()
    for text, is_keyword in parts:
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # Intentionally no bold, no underline for 'API endpoint' occurrences
        run.bold = False
        run.underline = False
    return para


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ===== TITLE PAGE =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('RESTful Services Integration Guide')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_para.add_run('Technical Reference Manual — Version 3.2')
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run('Prepared by: Platform Engineering Team\nLast Updated: February 2025')
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    toc_heading = doc.add_paragraph()
    toc_h_run = toc_heading.add_run('Table of Contents')
    toc_h_run.bold = True
    toc_h_run.font.size = Pt(16)
    toc_h_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    toc_items = [
        '1. Introduction to RESTful API Architecture .................. 3',
        '2. Authentication and Authorization .......................... 4',
        '3. Error Handling and Response Codes ......................... 5',
        '4. Best Practices and Rate Limiting .......................... 6',
        '5. Appendix: Glossary ........................................ 7',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_page_break()

    # ===== SECTION 1: Introduction to RESTful API Architecture =====
    # Heading
    s1_heading = doc.add_paragraph()
    s1_h_run = s1_heading.add_run('1. Introduction to RESTful API Architecture')
    s1_h_run.bold = True
    s1_h_run.font.size = Pt(16)
    s1_h_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Occurrence 1
    add_mixed_paragraph(doc, [
        ('This document provides a comprehensive reference for developers integrating with our platform services. Each ', False),
        ('API endpoint', True),
        (' in our system is designed following RESTful principles to ensure consistency and predictability across all service interactions.', False),
    ])

    doc.add_paragraph()

    # Subsection 1.1
    ss11 = doc.add_paragraph()
    ss11_run = ss11.add_run('1.1 Core Concepts')
    ss11_run.bold = True
    ss11_run.font.size = Pt(13)
    ss11_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    add_mixed_paragraph(doc, [
        ('REST (Representational State Transfer) is an architectural style that defines a set of constraints for creating web services. An ', False),
        ('API endpoint', True),
        (' represents a specific URL pattern that a client can call to perform a defined operation on a resource.', False),
    ])

    doc.add_paragraph()

    p_resources = doc.add_paragraph()
    r_resources = p_resources.add_run(
        'Resources in our platform are organized hierarchically under base URLs. Each resource type '
        'exposes a standard set of operations including retrieval, creation, update, and deletion. '
        'HTTP methods (GET, POST, PUT, DELETE, PATCH) map directly to these operations following '
        'established conventions that align with industry best practices.'
    )
    r_resources.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Subsection 1.2
    ss12 = doc.add_paragraph()
    ss12_run = ss12.add_run('1.2 Base URL Structure')
    ss12_run.bold = True
    ss12_run.font.size = Pt(13)
    ss12_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    p_url = doc.add_paragraph()
    r_url = p_url.add_run(
        'All service calls follow the pattern: https://api.platform.example.com/v3/{resource}. '
        'The versioning scheme ensures backward compatibility as the platform evolves. '
        'Query parameters can be appended to filter, sort, or paginate result sets.'
    )
    r_url.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Table: Common HTTP status codes
    t_heading = doc.add_paragraph()
    t_h_run = t_heading.add_run('Table 1: Standard HTTP Response Codes')
    t_h_run.bold = True
    t_h_run.font.size = Pt(11)
    t_h_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    headers_data = [('Status Code', 'Description'),
                    ('200 OK', 'Request succeeded; response body contains result'),
                    ('201 Created', 'Resource successfully created'),
                    ('400 Bad Request', 'Malformed request syntax or invalid parameters'),
                    ('401 Unauthorized', 'Authentication credentials missing or invalid')]
    for i, (col1, col2) in enumerate(headers_data):
        table.cell(i, 0).text = col1
        table.cell(i, 1).text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_page_break()

    # ===== SECTION 2: Authentication and Authorization =====
    s2_heading = doc.add_paragraph()
    s2_h_run = s2_heading.add_run('2. Authentication and Authorization')
    s2_h_run.bold = True
    s2_h_run.font.size = Pt(16)
    s2_h_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Occurrence 3
    add_mixed_paragraph(doc, [
        ('Secure access to the platform requires proper authentication credentials on every request. '
         'Each ', False),
        ('API endpoint', True),
        (' requires a valid bearer token issued through our OAuth 2.0 authorization service. '
         'Tokens expire after 3600 seconds and must be refreshed using the refresh token flow.', False),
    ])

    doc.add_paragraph()

    # Subsection 2.1
    ss21 = doc.add_paragraph()
    ss21_run = ss21.add_run('2.1 OAuth 2.0 Flow')
    ss21_run.bold = True
    ss21_run.font.size = Pt(13)
    ss21_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    p_oauth = doc.add_paragraph()
    r_oauth = p_oauth.add_run(
        'The authorization code flow is recommended for server-side applications. '
        'Client applications redirect the user to the authorization server, which issues a temporary '
        'authorization code. The application then exchanges this code for access and refresh tokens '
        'through a secure back-channel request.'
    )
    r_oauth.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Occurrence 4
    add_mixed_paragraph(doc, [
        ('For machine-to-machine integrations, the client credentials flow is preferred. The service account '
         'authenticates directly with the token endpoint, bypassing the user authorization step. Once obtained, '
         'the access token must be included as a Bearer token in the Authorization header of every call to any ', False),
        ('API endpoint', True),
        (' exposed by the platform.', False),
    ])

    doc.add_paragraph()

    # Subsection 2.2
    ss22 = doc.add_paragraph()
    ss22_run = ss22.add_run('2.2 Scopes and Permissions')
    ss22_run.bold = True
    ss22_run.font.size = Pt(13)
    ss22_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Occurrence 5
    add_mixed_paragraph(doc, [
        ('Access tokens include scope claims that restrict which operations a client can perform. '
         'The scope required for each ', False),
        ('API endpoint', True),
        (' is documented in the online API reference portal. Attempting to call an endpoint with insufficient '
         'scope results in a 403 Forbidden response.', False),
    ])

    doc.add_paragraph()

    p_scopes = doc.add_paragraph()
    r_scopes = p_scopes.add_run(
        'Available scopes include: read:users, write:users, read:orders, write:orders, admin:platform. '
        'Scopes are additive; a token may carry multiple scopes. Scope assignments are managed through '
        'the developer portal and reviewed quarterly by the security team.'
    )
    r_scopes.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_page_break()

    # ===== SECTION 3: Error Handling and Response Codes =====
    s3_heading = doc.add_paragraph()
    s3_h_run = s3_heading.add_run('3. Error Handling and Response Codes')
    s3_h_run.bold = True
    s3_h_run.font.size = Pt(16)
    s3_h_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Occurrence 6
    add_mixed_paragraph(doc, [
        ('Robust client applications must handle error responses gracefully. Every ', False),
        ('API endpoint', True),
        (' returns structured JSON error objects that provide machine-readable error codes alongside '
         'human-readable descriptions to facilitate debugging.', False),
    ])

    doc.add_paragraph()

    # Subsection 3.1
    ss31 = doc.add_paragraph()
    ss31_run = ss31.add_run('3.1 Standard Error Format')
    ss31_run.bold = True
    ss31_run.font.size = Pt(13)
    ss31_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    p_errfmt = doc.add_paragraph()
    r_errfmt = p_errfmt.add_run(
        'All error responses use a consistent envelope structure with the following fields: '
        '"error_code" (a machine-readable string identifier), "message" (a developer-facing explanation), '
        '"request_id" (a UUID for tracing), and "timestamp" (ISO 8601 format). '
        'Clients should log the request_id to assist platform support engineers in diagnosing issues.'
    )
    r_errfmt.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Subsection 3.2
    ss32 = doc.add_paragraph()
    ss32_run = ss32.add_run('3.2 Retry Strategy')
    ss32_run.bold = True
    ss32_run.font.size = Pt(13)
    ss32_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Occurrence 7
    add_mixed_paragraph(doc, [
        ('Transient failures (5xx status codes) are safe to retry with exponential backoff. '
         'The recommended initial backoff is 1 second, doubling on each retry up to a maximum of 32 seconds. '
         'Each ', False),
        ('API endpoint', True),
        (' includes a Retry-After header in rate-limit responses (HTTP 429) indicating '
         'when the client may safely resume requests.', False),
    ])

    doc.add_paragraph()

    p_idempotency = doc.add_paragraph()
    r_idempotency = p_idempotency.add_run(
        'For non-idempotent operations (POST, PATCH), include the Idempotency-Key header to prevent '
        'duplicate resource creation during retries. The platform caches idempotency keys for 24 hours. '
        'After this window, reusing the same key produces a new resource rather than returning the cached response.'
    )
    r_idempotency.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Subsection 3.3
    ss33 = doc.add_paragraph()
    ss33_run = ss33.add_run('3.3 Deprecation Notices')
    ss33_run.bold = True
    ss33_run.font.size = Pt(13)
    ss33_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    # Occurrence 8
    add_mixed_paragraph(doc, [
        ('When an ', False),
        ('API endpoint', True),
        (' is scheduled for removal, the platform sends deprecation notifications via the developer '
         'dashboard and includes a Deprecation header in response payloads at least 90 days before '
         'the endpoint is decommissioned. Clients should monitor these headers and migrate to successor '
         'endpoints promptly to avoid service disruption.', False),
    ])

    doc.add_paragraph()

    p_versioning = doc.add_paragraph()
    r_versioning = p_versioning.add_run(
        'Major version increments (v3 → v4) always introduce a minimum 12-month deprecation window '
        'for all affected endpoints. Minor version increments within the same major version are '
        'fully backward compatible and require no client-side changes.'
    )
    r_versioning.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_page_break()

    # ===== SECTION 4: Best Practices and Rate Limiting =====
    s4_heading = doc.add_paragraph()
    s4_h_run = s4_heading.add_run('4. Best Practices and Rate Limiting')
    s4_h_run.bold = True
    s4_h_run.font.size = Pt(16)
    s4_h_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    p_bp = doc.add_paragraph()
    r_bp = p_bp.add_run(
        'Rate limits are enforced per application credential set, using a sliding-window algorithm. '
        'Default limits are 1000 requests per minute for standard tiers and 10000 per minute for '
        'premium tiers. Contact platform support to request custom rate limit adjustments for '
        'high-volume production workloads.'
    )
    r_bp.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    p_caching = doc.add_paragraph()
    r_caching = p_caching.add_run(
        'Cache GET responses where appropriate using ETags and conditional requests. '
        'The platform sets cache-control headers on all cacheable resources. '
        'Respect these headers to reduce latency and avoid hitting rate limits unnecessarily.'
    )
    r_caching.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    p_pagination = doc.add_paragraph()
    r_pagination = p_pagination.add_run(
        'Always paginate large result sets using cursor-based pagination. The "next_cursor" field '
        'in list responses provides an opaque token to fetch the subsequent page. '
        'Avoid offset-based pagination for datasets exceeding 10,000 records.'
    )
    r_pagination.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_page_break()

    # ===== SECTION 5: Appendix =====
    s5_heading = doc.add_paragraph()
    s5_h_run = s5_heading.add_run('5. Appendix: Glossary')
    s5_h_run.bold = True
    s5_h_run.font.size = Pt(16)
    s5_h_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    glossary = [
        ('Access Token', 'A short-lived credential issued by the authorization server granting access to protected resources.'),
        ('API endpoint', 'A specific URL path that a RESTful service exposes to accept and process client requests for a given resource operation.'),
        ('Bearer Token', 'An opaque string passed in the Authorization header to authenticate API calls.'),
        ('CORS', 'Cross-Origin Resource Sharing — a browser mechanism that controls cross-origin HTTP requests.'),
        ('ETag', 'An identifier for a specific version of a resource, used for cache validation and conditional requests.'),
        ('Idempotency Key', 'A client-generated UUID included in request headers to enable safe retries of non-idempotent operations.'),
        ('OAuth 2.0', 'An industry-standard protocol for authorization, enabling delegated access without sharing credentials.'),
        ('Rate Limiting', 'A mechanism to control the rate of requests a client can make to prevent abuse and ensure fair use.'),
        ('Refresh Token', 'A long-lived token used to obtain new access tokens after the current one expires.'),
        ('REST', 'Representational State Transfer — an architectural style for distributed hypermedia systems.'),
        ('Scope', 'A parameter that limits the access rights granted by a token to specific operations or resources.'),
        ('Webhook', 'A user-defined HTTP callback triggered by specific platform events to notify external systems in real time.'),
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        r_term = p.add_run(f'{term}: ')
        r_term.bold = True
        r_term.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        r_def = p.add_run(definition)
        r_def.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
