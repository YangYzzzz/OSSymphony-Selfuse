"""
Initial Setup: Protected document with tracked changes
Task ID: writer_rm_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import hashlib
import base64
import struct
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_tracked_insertion(para, text, author, date, run_id):
    """Add a tracked insertion (w:ins) to a paragraph."""
    ins = para._element.makeelement(qn('w:ins'), {
        qn('w:id'): str(run_id),
        qn('w:author'): author,
        qn('w:date'): date,
    })
    r = ins.makeelement(qn('w:r'), {})
    rpr = r.makeelement(qn('w:rPr'), {})
    r.append(rpr)
    t = r.makeelement(qn('w:t'), {})
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    ins.append(r)
    para._element.append(ins)
    return ins


def add_tracked_deletion(para, text, author, date, run_id):
    """Add a tracked deletion (w:del) to a paragraph."""
    del_elem = para._element.makeelement(qn('w:del'), {
        qn('w:id'): str(run_id),
        qn('w:author'): author,
        qn('w:date'): date,
    })
    r = del_elem.makeelement(qn('w:r'), {})
    rpr = r.makeelement(qn('w:rPr'), {})
    r.append(rpr)
    dt = r.makeelement(qn('w:delText'), {})
    dt.text = text
    dt.set(qn('xml:space'), 'preserve')
    r.append(dt)
    del_elem.append(r)
    para._element.append(del_elem)
    return del_elem


def add_document_protection(doc, password='Secure123'):
    """Add document protection for tracked changes with password hash."""
    # Use legacy password hashing for .docx document protection
    # This uses the simple XOR-based hash that LibreOffice recognizes
    settings = doc.settings.element

    # Generate a simple hash for the password
    # Using the standard OOXML password hashing
    hash_val = 0
    for i, ch in enumerate(reversed(password)):
        char_val = ord(ch)
        # Rotate left by (i+1) bits within 15 bits
        intermediate = char_val << (i + 1)
        low15 = intermediate & 0x7FFF
        high = (intermediate >> 15) & 0x7FFF
        intermediate = low15 | high
        hash_val ^= intermediate
    hash_val ^= len(password)
    hash_val ^= 0xCE4B

    hash_hex = format(hash_val, '04X')

    protection = settings.makeelement(qn('w:documentProtection'), {
        qn('w:edit'): 'trackedChanges',
        qn('w:enforcement'): '1',
        qn('w:hash'): hash_hex,
        qn('w:cryptAlgorithmSid'): '1',
        qn('w:cryptSpinCount'): '0',
    })
    settings.append(protection)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============================================================
    # Title
    # ============================================================
    title = doc.add_heading('Quarterly Business Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    period = doc.add_paragraph()
    period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = period.add_run('Q4 2025 - October through December')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph('')  # spacer

    # ============================================================
    # Executive Summary
    # ============================================================
    doc.add_heading('Executive Summary', level=1)

    p1 = doc.add_paragraph(
        'Meridian Technologies experienced steady growth during the fourth quarter of 2025, '
        'with consolidated revenue reaching $47.3 million, representing a 12% increase over '
        'Q3 figures. Our enterprise software division continued to drive the majority of '
        'new client acquisitions, while the cloud infrastructure segment showed '
    )
    # Tracked change 1: insertion "remarkable" before "improvement"
    add_tracked_insertion(p1, 'remarkable ', 'Elena Rodriguez', '2025-12-18T09:15:00Z', 101)
    run = p1.add_run('improvement in margins.')
    run.font.size = Pt(11)

    p2 = doc.add_paragraph(
        'Key operational metrics indicate that customer retention rates remained above 94%, '
        'and our Net Promoter Score improved from 62 to 71 during this period. The '
        'leadership team has identified three strategic priorities for the upcoming fiscal year.'
    )

    # ============================================================
    # Revenue Analysis
    # ============================================================
    doc.add_heading('Revenue Analysis', level=1)

    p3 = doc.add_paragraph(
        'Total revenue for Q4 2025 was distributed across our four primary business units. '
        'The Enterprise Software division generated $18.9 million, representing 40% of total '
        'revenue. Cloud Infrastructure contributed $12.8 million (27%), Professional Services '
        'accounted for $9.5 million (20%), and the '
    )
    # Tracked change 2: deletion of "Legacy" before "Support"
    add_tracked_deletion(p3, 'Legacy ', 'James Park', '2025-12-19T14:30:00Z', 102)
    run = p3.add_run(
        'Support and Maintenance division brought in $6.1 million (13%).'
    )

    # Revenue table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Business Unit', 'Q4 Revenue ($M)', 'Q3 Revenue ($M)', 'Growth (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Enterprise Software', '$18.9', '$17.1', '10.5%'],
        ['Cloud Infrastructure', '$12.8', '$10.4', '23.1%'],
        ['Professional Services', '$9.5', '$9.2', '3.3%'],
        ['Support & Maintenance', '$6.1', '$5.5', '10.9%'],
        ['Total', '$47.3', '$42.2', '12.1%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')

    # Tracked change 3: insertion of a new analysis sentence
    p4 = doc.add_paragraph(
        'The Cloud Infrastructure segment showed the strongest quarter-over-quarter growth, '
        'driven primarily by new enterprise contracts signed in October and November.'
    )
    # Tracked change 4: insertion
    add_tracked_insertion(p4, ' The pipeline for Q1 2026 remains robust with $8.2M in committed deals.', 'Sarah Chen', '2025-12-20T10:45:00Z', 103)

    # ============================================================
    # Department Updates
    # ============================================================
    doc.add_heading('Department Updates', level=1)

    doc.add_heading('Engineering', level=2)
    p5 = doc.add_paragraph(
        'The engineering team completed the migration of our core platform to a '
        'microservices architecture, reducing average deployment time from 4 hours to '
    )
    # Tracked change 5: deletion "approximately "
    add_tracked_deletion(p5, 'approximately ', 'Marcus Johnson', '2025-12-17T16:20:00Z', 104)
    run = p5.add_run(
        '22 minutes. The team also shipped 3 major feature releases and resolved 847 '
        'customer-reported issues during the quarter.'
    )

    # Tracked change 6: insertion of new paragraph about hiring
    p6 = doc.add_paragraph('')
    add_tracked_insertion(p6, 'We successfully hired 12 senior engineers across the backend and infrastructure teams, bringing total engineering headcount to 156.', 'Elena Rodriguez', '2025-12-21T08:30:00Z', 105)

    doc.add_heading('Sales and Marketing', level=2)
    p7 = doc.add_paragraph(
        'The sales organization closed 47 new enterprise deals in Q4, with an average '
        'contract value of $285,000. Marketing campaigns generated over 3,200 qualified '
        'leads, a '
    )
    # Tracked change 7: insertion "significant "
    add_tracked_insertion(p7, 'significant ', 'James Park', '2025-12-19T11:00:00Z', 106)
    run = p7.add_run('increase from the 2,800 leads generated in Q3.')

    doc.add_heading('Human Resources', level=2)
    p8 = doc.add_paragraph(
        'Overall headcount grew from 412 to 438 employees during Q4. Employee satisfaction '
        'scores averaged 4.2 out of 5.0 on our internal survey. The voluntary turnover rate '
        'decreased to 8.3%, down from '
    )
    # Tracked change 8: deletion "an alarming "
    add_tracked_deletion(p8, 'an alarming ', 'Sarah Chen', '2025-12-22T13:15:00Z', 107)
    run = p8.add_run('11.7% in Q3.')

    # ============================================================
    # Financial Highlights
    # ============================================================
    doc.add_heading('Financial Highlights', level=1)

    p9 = doc.add_paragraph(
        'Operating expenses totaled $38.1 million in Q4, resulting in an operating margin '
        'of 19.4%. This represents a '
    )
    # Tracked change 9: insertion "meaningful "
    add_tracked_insertion(p9, 'meaningful ', 'Marcus Johnson', '2025-12-18T15:45:00Z', 108)
    run = p9.add_run(
        'improvement over the 16.8% margin achieved in Q3. Research and development '
        'spending was $11.2 million (23.7% of revenue), consistent with our long-term '
        'target range of 22-25%.'
    )

    p10 = doc.add_paragraph(
        'Cash and short-term investments stood at $89.4 million at quarter end, providing '
        'ample liquidity for planned strategic initiatives in 2026. '
    )
    # Tracked change 10: insertion about debt
    add_tracked_insertion(p10, 'Total outstanding debt was reduced by $3.5 million to $42.1 million.', 'Elena Rodriguez', '2025-12-23T09:00:00Z', 109)

    # ============================================================
    # Strategic Priorities for 2026
    # ============================================================
    doc.add_heading('Strategic Priorities for 2026', level=1)

    p11 = doc.add_paragraph(
        'The leadership team has established the following strategic priorities for the '
        'upcoming fiscal year:'
    )

    # Tracked change 11: insertion of "expanded " before "AI"
    item1 = doc.add_paragraph(style='List Number')
    run = item1.add_run('Accelerate ')
    add_tracked_insertion(item1, 'expanded ', 'James Park', '2025-12-20T14:00:00Z', 110)
    run = item1.add_run(
        'AI and machine learning capabilities across the product portfolio'
    )

    item2 = doc.add_paragraph(style='List Number')
    run = item2.add_run('Expand international presence with new offices in ')
    # Tracked change 15 (becomes our 14th): deletion of "London and " replaced by tracked insertion
    add_tracked_deletion(item2, 'London and ', 'Elena Rodriguez', '2025-12-20T15:00:00Z', 114)
    run = item2.add_run('Singapore')

    # Tracked change 12: deletion "tentatively " before "targeting"
    item3 = doc.add_paragraph(style='List Number')
    run = item3.add_run('Achieve $200 million annual recurring revenue, ')
    add_tracked_deletion(item3, 'tentatively ', 'Sarah Chen', '2025-12-21T16:30:00Z', 111)
    run = item3.add_run('targeting 18% year-over-year growth')

    # ============================================================
    # Risk Factors
    # ============================================================
    doc.add_heading('Risk Factors', level=1)

    p12 = doc.add_paragraph(
        'The company faces several risk factors that could impact performance in '
        'the coming quarters, including increased competition in the cloud infrastructure '
        'market, potential regulatory changes affecting data privacy requirements, and '
    )
    # Tracked change 13: insertion "ongoing global supply chain disruptions"
    add_tracked_insertion(p12, 'ongoing global supply chain disruptions affecting hardware procurement timelines', 'Marcus Johnson', '2025-12-22T10:30:00Z', 112)
    run = p12.add_run('.')

    # ============================================================
    # Conclusion
    # ============================================================
    doc.add_heading('Conclusion', level=1)

    p13 = doc.add_paragraph(
        'Q4 2025 demonstrated continued momentum across all business segments. The '
        'management team remains confident in our ability to execute on the 2026 strategic '
        'plan and deliver '
    )
    # Tracked change 14: insertion "sustainable long-term"
    add_tracked_insertion(p13, 'sustainable long-term ', 'Elena Rodriguez', '2025-12-23T11:00:00Z', 113)
    run = p13.add_run('value to our shareholders and customers.')

    doc.add_paragraph('')
    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = closing.add_run('Prepared by the Office of the CFO')
    run.italic = True
    run = closing.add_run('\nDecember 23, 2025')
    run.italic = True

    # ============================================================
    # Add document protection for tracked changes
    # ============================================================
    add_document_protection(doc, password='Secure123')

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
