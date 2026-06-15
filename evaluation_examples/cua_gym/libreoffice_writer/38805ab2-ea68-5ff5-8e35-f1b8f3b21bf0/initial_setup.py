"""
Initial Setup: Remove outer borders table task
Task ID: writer_tm_022
Domain: libreoffice_writer

Creates a document with a 3x5 table where all borders (inner and outer)
are 0.5pt solid black.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_022'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell using XML.
    Each border param is a dict like {'sz': '4', 'val': 'single', 'color': '000000'}
    or None to skip that border.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing borders element if any
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)

    borders_elem = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    border_map = {
        'top': top,
        'bottom': bottom,
        'left': left,
        'right': right,
    }

    for side, props in border_map.items():
        if props is not None:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} '
                f'w:val="{props["val"]}" '
                f'w:sz="{props["sz"]}" '
                f'w:space="0" '
                f'w:color="{props["color"]}"/>'
            )
            borders_elem.append(border)

    tcPr.append(borders_elem)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Modern Layout Report', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add intro paragraph
    intro = doc.add_paragraph(
        'The following table summarizes quarterly performance metrics '
        'across our three regional divisions for the fiscal year 2025.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Create 5 rows x 3 cols table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # Table data - realistic content
    data = [
        ['Region', 'Q1 Revenue ($K)', 'Q2 Revenue ($K)'],
        ['North America', '2,345', '2,580'],
        ['Europe & UK', '1,890', '1,975'],
        ['Asia-Pacific', '3,120', '3,410'],
        ['Latin America', '987', '1,045'],
    ]

    # Populate table
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''  # clear default
            para = cell.paragraphs[0]
            run = para.add_run(val)
            if i == 0:
                # Header row - bold
                run.bold = True
                run.font.size = Pt(11)
            else:
                run.font.size = Pt(11)

    # Set all borders to 0.5pt (4 eighths of a point) solid black on every cell
    border_props = {'sz': '4', 'val': 'single', 'color': '000000'}

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            set_cell_borders(
                cell,
                top=border_props,
                bottom=border_props,
                left=border_props,
                right=border_props,
            )

    # Add a closing paragraph
    closing = doc.add_paragraph(
        'Data compiled by the Finance Department. All figures are preliminary '
        'and subject to final audit review.'
    )
    closing.paragraph_format.space_before = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
