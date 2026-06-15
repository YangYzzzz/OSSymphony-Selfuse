"""
Reward Script: Remove outer borders of table, keep inner gridlines
Task ID: writer_tm_022
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Top row outer border (top edge) removed
  Component 2 (0.35): Bottom row outer border (bottom edge) removed
  Component 3 (0.15): Left column outer border (left edge) removed
  Component 4 (0.15): Right column outer border (right edge) removed
  Bonus validation: Inner borders still present (gates full score)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_022'


def get_cell_border(cell, border_name):
    """
    Get border info for a cell. border_name is one of: top, bottom, left, right.
    Returns (val, sz) tuple or (None, None) if not found.
    """
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return (None, None)
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return (None, None)
    border_el = tcBorders.find(qn('w:' + border_name))
    if border_el is None:
        return (None, None)
    val = border_el.get(qn('w:val'))
    sz = border_el.get(qn('w:sz'))
    return (val, sz)


def is_border_removed(cell, border_name):
    """Check if a specific border on a cell is removed (none/nil/sz=0)."""
    val, sz = get_cell_border(cell, border_name)
    if val in ('none', 'nil'):
        return True
    if val is None:
        # No explicit border set; inherited from style (likely single)
        return False
    if sz is not None and int(sz) == 0:
        return True
    return False


def is_border_present(cell, border_name):
    """Check if a specific border on a cell is present (single with sz > 0)."""
    val, sz = get_cell_border(cell, border_name)
    if val == 'single' and sz is not None and int(sz) > 0:
        return True
    if val is None:
        # Inherited from TableGrid style — borders are present by default
        return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file {}: {}".format(file_path, e))
        print("REWARD: 0.0")
        return 0.0

    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    print("Table: {} rows x {} cols".format(num_rows, num_cols))

    if num_rows < 2 or num_cols < 2:
        print("CRITICAL: Table too small for meaningful border check")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Top row outer border removed (0.35 points)
    # All cells in row 0 should have top=none
    try:
        top_removed_count = 0
        for ci in range(num_cols):
            cell = table.cell(0, ci)
            if is_border_removed(cell, 'top'):
                top_removed_count += 1
        if top_removed_count == num_cols:
            print("PASS: Component 1 - All top-row top borders removed ({}/{}) (0.35 pts)".format(
                top_removed_count, num_cols))
            total_score += 0.35
        else:
            print("FAIL: Component 1 - Top-row top borders removed: {}/{} (expected all)".format(
                top_removed_count, num_cols))
    except Exception as e:
        print("ERROR: Component 1 - {}".format(e))

    # Component 2: Bottom row outer border removed (0.35 points)
    # All cells in last row should have bottom=none
    try:
        bottom_removed_count = 0
        last_row = num_rows - 1
        for ci in range(num_cols):
            cell = table.cell(last_row, ci)
            if is_border_removed(cell, 'bottom'):
                bottom_removed_count += 1
        if bottom_removed_count == num_cols:
            print("PASS: Component 2 - All bottom-row bottom borders removed ({}/{}) (0.35 pts)".format(
                bottom_removed_count, num_cols))
            total_score += 0.35
        else:
            print("FAIL: Component 2 - Bottom-row bottom borders removed: {}/{} (expected all)".format(
                bottom_removed_count, num_cols))
    except Exception as e:
        print("ERROR: Component 2 - {}".format(e))

    # Component 3: Left column outer border removed (0.15 points)
    # All cells in column 0 should have left=none
    try:
        left_removed_count = 0
        for ri in range(num_rows):
            cell = table.cell(ri, 0)
            if is_border_removed(cell, 'left'):
                left_removed_count += 1
        if left_removed_count == num_rows:
            print("PASS: Component 3 - All left-col left borders removed ({}/{}) (0.15 pts)".format(
                left_removed_count, num_rows))
            total_score += 0.15
        else:
            print("FAIL: Component 3 - Left-col left borders removed: {}/{} (expected all)".format(
                left_removed_count, num_rows))
    except Exception as e:
        print("ERROR: Component 3 - {}".format(e))

    # Component 4: Right column outer border removed (0.15 points)
    # All cells in last column should have right=none
    try:
        right_removed_count = 0
        last_col = num_cols - 1
        for ri in range(num_rows):
            cell = table.cell(ri, last_col)
            if is_border_removed(cell, 'right'):
                right_removed_count += 1
        if right_removed_count == num_rows:
            print("PASS: Component 4 - All right-col right borders removed ({}/{}) (0.15 pts)".format(
                right_removed_count, num_rows))
            total_score += 0.15
        else:
            print("FAIL: Component 4 - Right-col right borders removed: {}/{} (expected all)".format(
                right_removed_count, num_rows))
    except Exception as e:
        print("ERROR: Component 4 - {}".format(e))

    # Validation gate: Inner borders must still be present
    # If inner borders are also removed, reduce score (task only asked to remove outer)
    # Check a sample of inner borders: horizontal between rows, vertical between cols
    try:
        inner_present = 0
        inner_total = 0

        # Check horizontal inner borders: bottom of row 0..N-2 (interior rows)
        for ri in range(num_rows - 1):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                if is_border_present(cell, 'bottom'):
                    inner_present += 1
                inner_total += 1

        # Check vertical inner borders: right of col 0..N-2
        for ri in range(num_rows):
            for ci in range(num_cols - 1):
                cell = table.cell(ri, ci)
                if is_border_present(cell, 'right'):
                    inner_present += 1
                inner_total += 1

        inner_ratio = inner_present / inner_total if inner_total > 0 else 0
        print("Inner borders present: {}/{} ({:.0%})".format(inner_present, inner_total, inner_ratio))
        if inner_ratio < 0.8:
            # Inner borders were also removed - penalize
            penalty = total_score * 0.5
            print("PENALTY: Inner borders largely removed - reducing score by {:.2f}".format(penalty))
            total_score -= penalty
    except Exception as e:
        print("ERROR: Inner border check - {}".format(e))

    final_score = max(0.0, min(total_score, 1.0))
    print("")
    print("Score: {}/1.0".format(round(total_score, 2)))
    print("REWARD: {}".format(round(final_score, 2)))
    return final_score


# Default: test against canonical artifact path
file_path = WORKDIR + '/' + TASK_ID + '.docx'
if not os.path.exists(file_path):
    print("File not found: " + file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
