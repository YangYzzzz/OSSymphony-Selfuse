"""
Initial Setup: Company Handbook with Heading 1 and Heading 2 paragraphs (no numbering)
Task ID: writer_list_029
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'handbook'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default empty paragraph if present
    # (We'll build the document from scratch with proper structure)

    # Chapter 1: Employment Policies
    h1 = doc.add_paragraph('Employment Policies', style='Heading 1')

    h2 = doc.add_paragraph('Equal Opportunity', style='Heading 2')
    p = doc.add_paragraph(
        'Our company is committed to providing equal employment opportunities to all employees '
        'and applicants. We do not discriminate on the basis of race, color, national origin, '
        'gender, age, disability, religion, or any other characteristic protected by law.'
    )

    h2 = doc.add_paragraph('Code of Conduct', style='Heading 2')
    p = doc.add_paragraph(
        'All employees are expected to maintain the highest standards of professional conduct. '
        'This includes treating colleagues, clients, and partners with respect and integrity, '
        'and adhering to all company policies and applicable laws at all times.'
    )

    h2 = doc.add_paragraph('Attendance Requirements', style='Heading 2')
    p = doc.add_paragraph(
        'Regular and punctual attendance is essential to the smooth operation of our business. '
        'Employees are expected to report to work as scheduled. Any absences must be communicated '
        'to the direct supervisor at least one hour before the start of the workday.'
    )

    # Chapter 2: Compensation and Benefits
    h1 = doc.add_paragraph('Compensation and Benefits', style='Heading 1')

    h2 = doc.add_paragraph('Salary Structure', style='Heading 2')
    p = doc.add_paragraph(
        'The company maintains a competitive salary structure based on industry benchmarks, '
        'employee performance, and years of experience. Salary reviews are conducted annually '
        'and may result in merit-based increases. The compensation philosophy aims to attract '
        'and retain top talent.'
    )

    h2 = doc.add_paragraph('Health Insurance', style='Heading 2')
    p = doc.add_paragraph(
        'Full-time employees are eligible to enroll in the company health insurance plan after '
        'completing 30 days of employment. The company covers 80% of the premium for the employee '
        'and offers options for dependents coverage. Open enrollment occurs each November for '
        'coverage beginning the following January.'
    )

    # Chapter 3: Safety and Security
    h1 = doc.add_paragraph('Safety and Security', style='Heading 1')

    h2 = doc.add_paragraph('Emergency Procedures', style='Heading 2')
    p = doc.add_paragraph(
        'In the event of an emergency, employees should follow the posted evacuation routes and '
        'assemble at the designated meeting point in the parking lot. Fire drills are conducted '
        'twice yearly. Emergency contact numbers are posted in each department and on the company '
        'intranet.'
    )

    h2 = doc.add_paragraph('Workplace Safety', style='Heading 2')
    p = doc.add_paragraph(
        'The company is committed to maintaining a safe and healthy work environment. Employees '
        'must report any unsafe conditions or accidents to the Safety Officer immediately. '
        'Personal protective equipment is provided where required by safety regulations. '
        'Safety training is mandatory for all new hires during their first week.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
