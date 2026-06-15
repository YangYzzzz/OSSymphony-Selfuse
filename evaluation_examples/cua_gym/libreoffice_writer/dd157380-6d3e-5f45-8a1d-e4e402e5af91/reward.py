"""
Reward Script: Multi-level heading numbering for company handbook
Task ID: writer_list_029
Domain: libreoffice_writer
Scoring:
  Component 1: Heading 1 paragraphs have numbering applied (numPr) - 0.3 pts
  Component 2: Heading 2 paragraphs have numbering applied (numPr) - 0.3 pts
  Component 3: Multi-level numbering uses correct formats
               (Level 0 decimal '%1.' for H1, Level 1 parent-child '%1.%2.' for H2) - 0.4 pts
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_list_029'

# Expected heading text structure
H1_HEADINGS = ['Employment Policies', 'Compensation and Benefits', 'Safety and Security']
H2_HEADINGS = ['Equal Opportunity', 'Code of Conduct', 'Attendance Requirements',
               'Salary Structure', 'Health Insurance', 'Emergency Procedures', 'Workplace Safety']


def get_numPr(para):
    """Extract numId and ilvl from paragraph's numPr element. Returns (numId, ilvl) or (None, None)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None, None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(qn('w:ilvl'))
    numId_el = numPr.find(qn('w:numId'))
    ilvl_val = int(ilvl_el.get(qn('w:val'))) if ilvl_el is not None else None
    numId_val = int(numId_el.get(qn('w:val'))) if numId_el is not None else None
    # numId=0 means numbering is explicitly turned off
    if numId_val == 0:
        return None, None
    return numId_val, ilvl_val


def get_abstractNum_for_numId(numbering_part, numId):
    """Resolve numId -> abstractNumId using the numbering part."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for num_el in numbering_part._element.findall(qn('w:num')):
        if int(num_el.get(qn('w:numId'))) == numId:
            abst_ref = num_el.find(qn('w:abstractNumId'))
            if abst_ref is not None:
                return int(abst_ref.get(qn('w:val')))
    return None


def get_abstractNum_levels(numbering_part, abstractNumId):
    """Get level definitions from an abstractNum. Returns dict: ilvl -> {numFmt, lvlText, multiLevelType}."""
    levels = {}
    for abst_el in numbering_part._element.findall(qn('w:abstractNum')):
        if int(abst_el.get(qn('w:abstractNumId'))) == abstractNumId:
            mlt_el = abst_el.find(qn('w:multiLevelType'))
            ml_type = mlt_el.get(qn('w:val')) if mlt_el is not None else 'unknown'
            for lvl_el in abst_el.findall(qn('w:lvl')):
                ilvl = int(lvl_el.get(qn('w:ilvl')))
                numFmt_el = lvl_el.find(qn('w:numFmt'))
                lvlText_el = lvl_el.find(qn('w:lvlText'))
                numFmt = numFmt_el.get(qn('w:val')) if numFmt_el is not None else None
                lvlText = lvlText_el.get(qn('w:val')) if lvlText_el is not None else None
                levels[ilvl] = {'numFmt': numFmt, 'lvlText': lvlText, 'multiLevelType': ml_type}
    return levels


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect heading paragraphs
    h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
    h2_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 2']

    # Precondition: verify expected headings exist
    h1_texts = [p.text for p in h1_paras]
    h2_texts = [p.text for p in h2_paras]
    print(f"INFO: Found {len(h1_paras)} Heading 1 paragraphs: {h1_texts}")
    print(f"INFO: Found {len(h2_paras)} Heading 2 paragraphs: {h2_texts}")

    # Get numbering part for later use
    numbering_part = None
    if hasattr(doc.part, 'numbering_part') and doc.part.numbering_part is not None:
        numbering_part = doc.part.numbering_part

    # -------------------------------------------------------------------------
    # Component 1: Heading 1 paragraphs have numbering applied (0.3 points)
    # Initial env: Heading 1 paragraphs have NO numPr
    # Golden env:  Heading 1 paragraphs have numId>0, ilvl=0
    # -------------------------------------------------------------------------
    try:
        if len(h1_paras) == 0:
            print("FAIL: Component 1 — No Heading 1 paragraphs found")
        else:
            h1_numbered = []
            h1_not_numbered = []
            for p in h1_paras:
                numId_val, ilvl_val = get_numPr(p)
                if numId_val is not None and ilvl_val == 0:
                    h1_numbered.append(p.text)
                else:
                    h1_not_numbered.append(p.text)

            if len(h1_numbered) == len(h1_paras) and len(h1_paras) >= 3:
                print(f"PASS: Component 1 — All {len(h1_numbered)} Heading 1 paragraphs have numbering at ilvl=0 (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — Only {len(h1_numbered)}/{len(h1_paras)} Heading 1 paragraphs numbered. "
                      f"Missing: {h1_not_numbered}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Heading 2 paragraphs have numbering applied (0.3 points)
    # Initial env: Heading 2 paragraphs have NO numPr
    # Golden env:  Heading 2 paragraphs have numId>0, ilvl=1
    # -------------------------------------------------------------------------
    try:
        if len(h2_paras) == 0:
            print("FAIL: Component 2 — No Heading 2 paragraphs found")
        else:
            h2_numbered = []
            h2_not_numbered = []
            for p in h2_paras:
                numId_val, ilvl_val = get_numPr(p)
                if numId_val is not None and ilvl_val == 1:
                    h2_numbered.append(p.text)
                else:
                    h2_not_numbered.append(p.text)

            if len(h2_numbered) == len(h2_paras) and len(h2_paras) >= 5:
                print(f"PASS: Component 2 — All {len(h2_numbered)} Heading 2 paragraphs have numbering at ilvl=1 (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Only {len(h2_numbered)}/{len(h2_paras)} Heading 2 paragraphs numbered at ilvl=1. "
                      f"Missing/wrong: {h2_not_numbered}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Multi-level numbering format — Level 0 decimal, Level 1 parent-child (0.4 points)
    # Task requires:
    #   Level 1 (ilvl=0): Arabic numerals (1, 2, 3) → numFmt=decimal
    #   Level 2 (ilvl=1): parent-child numbering (1.1, 1.2) → lvlText contains %1 and %2
    # Both H1 and H2 must reference the same numId and that numId must be multilevel
    # -------------------------------------------------------------------------
    try:
        if numbering_part is None:
            print("FAIL: Component 3 — No numbering part found in document")
        else:
            # Get the numId used by headings
            h1_numIds = set()
            h2_numIds = set()
            for p in h1_paras:
                numId_val, _ = get_numPr(p)
                if numId_val is not None:
                    h1_numIds.add(numId_val)
            for p in h2_paras:
                numId_val, _ = get_numPr(p)
                if numId_val is not None:
                    h2_numIds.add(numId_val)

            # H1 and H2 should share the same numId (same list)
            shared_numIds = h1_numIds & h2_numIds
            if not shared_numIds:
                print(f"FAIL: Component 3 — Heading 1 (numIds={h1_numIds}) and Heading 2 (numIds={h2_numIds}) "
                      f"do not share a common numId; they must belong to the same multi-level list")
            else:
                # Check the abstract numbering definition
                numId = next(iter(shared_numIds))
                abstractNumId = get_abstractNum_for_numId(numbering_part, numId)
                if abstractNumId is None:
                    print(f"FAIL: Component 3 — Cannot resolve abstractNumId for numId={numId}")
                else:
                    levels = get_abstractNum_levels(numbering_part, abstractNumId)
                    print(f"INFO: abstractNumId={abstractNumId}, levels={levels}")

                    ml_type = levels.get(0, {}).get('multiLevelType', 'unknown') if levels else 'unknown'
                    lvl0 = levels.get(0, {})
                    lvl1 = levels.get(1, {})

                    lvl0_numFmt = lvl0.get('numFmt')
                    lvl0_lvlText = lvl0.get('lvlText', '')
                    lvl1_numFmt = lvl1.get('numFmt')
                    lvl1_lvlText = lvl1.get('lvlText', '')

                    # Level 0: must be decimal (Arabic numerals)
                    lvl0_ok = (lvl0_numFmt == 'decimal')
                    # Level 1: must be decimal AND lvlText must include both parent (%1) and child (%2) counters
                    lvl1_ok = (lvl1_numFmt == 'decimal' and
                               lvl1_lvlText is not None and
                               '%1' in lvl1_lvlText and
                               '%2' in lvl1_lvlText)
                    # Must be a multilevel list (not singleLevel)
                    is_multilevel = (ml_type == 'multilevel') or (1 in levels)

                    comp3_pass = lvl0_ok and lvl1_ok and is_multilevel
                    if comp3_pass:
                        total_score += 0.4
                        print(f"PASS: Component 3 — Multi-level numbering correct: "
                              f"L0 numFmt={lvl0_numFmt!r} lvlText={lvl0_lvlText!r}, "
                              f"L1 numFmt={lvl1_numFmt!r} lvlText={lvl1_lvlText!r} (0.4 pts)")
                    else:
                        details = []
                        if not lvl0_ok:
                            details.append(f"Level 0 numFmt={lvl0_numFmt!r} (expected 'decimal')")
                        if not lvl1_ok:
                            details.append(f"Level 1 numFmt={lvl1_numFmt!r} lvlText={lvl1_lvlText!r} "
                                           f"(expected decimal with '%1...%2' pattern)")
                        if not is_multilevel:
                            details.append(f"multiLevelType={ml_type!r} (expected 'multilevel' or 2+ levels)")
                        print(f"FAIL: Component 3 — {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/Desktop/handbook.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
