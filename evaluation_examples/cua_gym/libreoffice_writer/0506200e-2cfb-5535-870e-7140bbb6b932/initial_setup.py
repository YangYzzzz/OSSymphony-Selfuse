"""
Initial Setup: Create a Product_List document with a 3x10 table.
All cells have default top vertical alignment. Row heights vary.
Task ID: writer_tm_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add title
    heading = doc.add_heading('Product Inventory List', level=1)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('Updated: March 2026 | Warehouse B-7 Catalog')

    # Create 3-column, 11-row table (1 header + 10 data rows)
    table = doc.add_table(rows=11, cols=3)
    table.style = 'Table Grid'

    # Header row
    headers = ['Product Name', 'Category', 'Unit Price (USD)']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Data rows - realistic product data
    data = [
        ['Ergonomic Office Chair',     'Furniture',     '$349.99'],
        ['Wireless Bluetooth Mouse',   'Electronics',   '$29.50'],
        ['Standing Desk Converter',    'Furniture',     '$189.00'],
        ['USB-C Docking Station',      'Electronics',   '$124.95'],
        ['Noise-Canceling Headphones', 'Electronics',   '$275.00'],
        ['LED Desk Lamp',              'Lighting',      '$45.80'],
        ['Mechanical Keyboard',        'Electronics',   '$89.99'],
        ['Monitor Arm Mount',          'Accessories',   '$67.50'],
        ['Webcam HD 1080p',            'Electronics',   '$54.25'],
        ['Cable Management Kit',       'Accessories',   '$18.75'],
    ]

    # Varying row heights (in cm): 1.0 to 2.5
    row_heights_cm = [1.0, 1.5, 2.0, 1.2, 2.5, 1.8, 1.3, 2.2, 1.6, 2.0]

    for i, (row_data, height_cm) in enumerate(zip(data, row_heights_cm), start=1):
        row = table.rows[i]
        # Set row height
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = trPr.makeelement(qn('w:trHeight'), {
            qn('w:val'): str(int(height_cm * 567)),  # 1 cm ~ 567 twips
            qn('w:hRule'): 'atLeast'
        })
        trPr.append(trHeight)

        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val

    # Ensure all cells have default (top) vertical alignment - this is the default,
    # but let's be explicit by NOT setting any vAlign (top is default when absent)

    doc.add_paragraph('')  # spacer
    doc.add_paragraph('Note: All prices are subject to quarterly review.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
