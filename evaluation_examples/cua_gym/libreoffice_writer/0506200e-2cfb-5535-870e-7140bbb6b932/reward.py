"""
Reward Script: Set vertical alignment of all table cells to Center
Task ID: writer_tm_043
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): At least 80% of cells have center vertical alignment
  Component 2 (0.3): ALL cells have center vertical alignment (100%)
  Component 3 (0.3): Table content is preserved (text not corrupted)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_043'


def persist_app_state(domain: str):
    """Best-effort save via Ctrl+S before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_cell_valign(cell):
    """Get vertical alignment of a table cell. Returns string or None."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        va = tcPr.find(qn('w:vAlign'))
        if va is not None:
            return va.get(qn('w:val'))
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    total_cells = 0
    center_count = 0

    for row in table.rows:
        for cell in row.cells:
            total_cells += 1
            valign = get_cell_valign(cell)
            if valign == 'center':
                center_count += 1

    if total_cells == 0:
        print("FAIL: Table has no cells")
        print("REWARD: 0.0")
        return 0.0

    center_ratio = center_count / total_cells
    print(f"INFO: {center_count}/{total_cells} cells have center vertical alignment ({center_ratio:.1%})")

    # Component 1: At least 80% of cells are center-aligned (0.4 points)
    # This checks that the bulk of the work is done
    try:
        if center_ratio >= 0.8:
            print(f"PASS: Component 1 -- >= 80% cells center-aligned ({center_ratio:.1%}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 -- only {center_ratio:.1%} cells center-aligned, need >= 80%")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: ALL cells are center-aligned (0.3 points)
    # Full marks only when every single cell is done
    try:
        if center_count == total_cells:
            print(f"PASS: Component 2 -- all {total_cells} cells center-aligned (0.3 pts)")
            total_score += 0.3
        else:
            missing = total_cells - center_count
            print(f"FAIL: Component 2 -- {missing} cells still not center-aligned")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Table content preserved -- check that cell text is not empty/corrupted (0.3 points)
    # This verifies the task didn't destroy data. Only scores if center alignment was actually applied
    # (anchored to the task change: alignment changed AND content intact)
    try:
        # Expected: 3x11 table (header + 10 data rows, 3 columns)
        expected_rows = 11
        expected_cols = 3
        row_count = len(table.rows)
        col_count = len(table.columns)

        structure_ok = (row_count == expected_rows and col_count == expected_cols)

        # Check header row text is intact
        header_texts = [table.cell(0, c).text.strip() for c in range(min(col_count, 3))]
        expected_headers = ['Product Name', 'Category', 'Unit Price (USD)']
        headers_ok = (header_texts == expected_headers)

        # Non-empty data cells (at least 80% should have text)
        non_empty = 0
        data_cells = 0
        for i in range(1, row_count):
            for j in range(col_count):
                data_cells += 1
                if table.cell(i, j).text.strip():
                    non_empty += 1
        data_ok = data_cells > 0 and (non_empty / data_cells) >= 0.8

        # Only award if center alignment is actually applied (anchored to task change)
        if center_ratio >= 0.8 and structure_ok and headers_ok and data_ok:
            print(f"PASS: Component 3 -- table structure ({row_count}x{col_count}) and content preserved with alignment applied (0.3 pts)")
            total_score += 0.3
        else:
            if center_ratio < 0.8:
                print(f"FAIL: Component 3 -- center alignment not sufficiently applied (prerequisite)")
            if not structure_ok:
                print(f"FAIL: Component 3 -- table structure changed: {row_count}x{col_count} vs expected {expected_rows}x{expected_cols}")
            if not headers_ok:
                print(f"FAIL: Component 3 -- headers changed: {header_texts} vs expected {expected_headers}")
            if not data_ok:
                print(f"FAIL: Component 3 -- too many empty data cells: {non_empty}/{data_cells}")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
