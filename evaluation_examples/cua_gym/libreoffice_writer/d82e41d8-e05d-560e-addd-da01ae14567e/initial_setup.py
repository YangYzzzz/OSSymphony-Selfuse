"""
Initial Setup: Create a Writer document with heading structure for outline demotion task.
Task ID: writer_fp_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_040'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Style setup ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # =====================================================
    # Chapter 1: Introduction (Heading 1)
    # =====================================================
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'This report presents the findings of a comprehensive twelve-month '
        'study conducted by the Meridian Research Institute examining the '
        'impact of remote work policies on employee productivity and '
        'organizational performance across multiple industries.'
    )
    doc.add_paragraph(
        'The research was motivated by the rapid shift toward distributed '
        'work arrangements following global disruptions in 2020. Despite '
        'widespread adoption, empirical evidence regarding long-term effects '
        'on both individual output and team cohesion remained limited at the '
        'time of study inception.'
    )
    doc.add_paragraph(
        'Our investigation spanned 47 organizations across the technology, '
        'financial services, and healthcare sectors, encompassing over 3,200 '
        'individual participants. Data collection methods included quarterly '
        'surveys, biometric productivity tracking, and structured interviews '
        'with department leads.'
    )

    # =====================================================
    # Chapter 2: Methods (Heading 1)
    # =====================================================
    doc.add_heading('Methods', level=1)

    doc.add_paragraph(
        'A mixed-methods approach was employed, combining quantitative '
        'performance metrics with qualitative assessments. Participants were '
        'randomly assigned to one of three cohorts: fully remote, hybrid '
        '(three days in-office), or fully on-site, each containing '
        'approximately 1,070 individuals.'
    )
    doc.add_paragraph(
        'Quantitative data included weekly output measurements normalized '
        'by role category, error rates in deliverables, and communication '
        'frequency metrics derived from collaboration platform logs. '
        'Qualitative interviews were transcribed and coded using grounded '
        'theory techniques by two independent reviewers.'
    )
    doc.add_paragraph(
        'Statistical analysis was performed using a linear mixed-effects '
        'model with random intercepts for organization and fixed effects '
        'for cohort assignment, tenure, and job category. Significance '
        'thresholds were set at alpha = 0.05 with Bonferroni corrections '
        'applied for multiple comparisons.'
    )

    # =====================================================
    # Chapter 3: Results (Heading 1)
    # =====================================================
    doc.add_heading('Results', level=1)

    doc.add_paragraph(
        'Across all sectors, the hybrid cohort demonstrated the highest '
        'mean productivity index (87.3 on a normalized 100-point scale), '
        'followed by fully remote (82.1) and on-site (79.6). The '
        'difference between hybrid and on-site cohorts was statistically '
        'significant (p < 0.001), while the remote-versus-on-site '
        'comparison approached but did not reach significance (p = 0.054).'
    )
    doc.add_paragraph(
        'Sector-level analysis revealed notable heterogeneity. Technology '
        'firms showed the largest productivity gains under remote '
        'arrangements (mean delta = +9.2 points), whereas healthcare '
        'organizations exhibited a slight decline (-1.8 points) attributed '
        'to the inherently collaborative nature of clinical workflows.'
    )

    # =====================================================
    # Chapter 4: Summary of Findings (Heading 1)
    #   4.1 Key Metrics (Heading 2)
    #   4.2 Statistical Significance (Heading 2)
    # =====================================================
    doc.add_heading('Summary of Findings', level=1)

    doc.add_paragraph(
        'The overall findings indicate a clear advantage for hybrid work '
        'models across most organizational contexts. This section '
        'consolidates the primary metrics and their statistical reliability '
        'for reference by decision-makers and future researchers.'
    )

    # Sub-heading: Key Metrics (Heading 2)
    doc.add_heading('Key Metrics', level=2)

    doc.add_paragraph(
        'The three principal metrics tracked throughout the study were '
        'the Normalized Productivity Index (NPI), the Deliverable Error '
        'Rate (DER), and the Collaboration Frequency Score (CFS). The '
        'hybrid cohort scored highest on NPI (87.3), lowest on DER (2.4%), '
        'and maintained a CFS comparable to on-site workers (71.8 versus '
        '73.2).'
    )
    doc.add_paragraph(
        'Retention rates were also markedly higher in the hybrid group '
        '(94.1% annualized) compared to on-site (88.7%) and fully remote '
        '(90.3%). Exit interviews cited schedule flexibility and reduced '
        'commute burden as the two most frequently mentioned factors '
        'influencing the decision to stay.'
    )

    # Sub-heading: Statistical Significance (Heading 2)
    doc.add_heading('Statistical Significance', level=2)

    doc.add_paragraph(
        'After applying Bonferroni corrections, the hybrid-versus-on-site '
        'comparison remained highly significant across all three metrics '
        '(NPI: p < 0.001; DER: p = 0.003; CFS: p = 0.041). The '
        'remote-versus-on-site comparison reached significance only for '
        'NPI (p = 0.012) and DER (p = 0.028), while CFS did not differ '
        'meaningfully (p = 0.312).'
    )
    doc.add_paragraph(
        'Effect sizes, measured using Cohen\'s d, ranged from medium '
        '(d = 0.45 for CFS) to large (d = 0.81 for NPI) when comparing '
        'hybrid and on-site cohorts. These effect sizes suggest practical '
        'significance in addition to statistical significance, reinforcing '
        'the business case for hybrid arrangements.'
    )

    # =====================================================
    # Chapter 5: Conclusion (Heading 1)
    # =====================================================
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        'This study provides robust evidence that hybrid work models '
        'outperform both fully remote and fully on-site arrangements '
        'across the majority of measured dimensions. Organizations seeking '
        'to optimize productivity while maintaining employee satisfaction '
        'should consider structured hybrid policies with clearly defined '
        'in-office collaboration days.'
    )
    doc.add_paragraph(
        'Future research should extend the observation period beyond '
        'twelve months, incorporate additional sectors such as education '
        'and manufacturing, and examine potential confounders including '
        'managerial experience with distributed teams and the availability '
        'of purpose-built remote work infrastructure.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
