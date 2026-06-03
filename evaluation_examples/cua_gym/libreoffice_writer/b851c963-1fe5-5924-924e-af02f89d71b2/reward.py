"""
Reward Script: SRS Document Creation in LibreOffice Writer
Task ID: writer_wf_020
Domain: libreoffice_writer
Scoring:
  C1: Title text matches (0.10)
  C2: TOC section exists with Heading 1 (0.10)
  C3: Six required Heading 1 sections present (0.20)
  C4: Functional Requirements table - 4 cols, 7 rows, correct headers (0.20)
  C5: Non-Functional Requirements - 4 bulleted list items (0.15)
  C6: Glossary table - 2 cols, 6 rows, correct headers (0.15)
  C7: Document Overview & System Description sections have body text (0.10)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_020'


def verify_task(file_path):
    """
    Verify SRS document creation with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    paragraphs = doc.paragraphs
    tables = doc.tables

    # If document is essentially empty, return 0 immediately
    if len(paragraphs) == 0 and len(tables) == 0:
        print("FAIL: Document is empty (no paragraphs or tables)")
        print("REWARD: 0.0")
        return 0.0

    # Build helper structures
    heading1_texts = []
    all_styles = []
    for p in paragraphs:
        style_name = p.style.name if p.style else ''
        all_styles.append((style_name, p.text.strip()))
        if style_name == 'Heading 1':
            heading1_texts.append(p.text.strip().lower())

    # Component 1: Title text matches (0.10 points)
    try:
        title_found = False
        for p in paragraphs:
            style_name = p.style.name if p.style else ''
            text = p.text.strip().lower()
            if 'srs' in text and 'inventory management system' in text and 'v2.0' in text:
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 - Title contains SRS Inventory Management System v2.0 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 - Title not found with expected text")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: TOC section exists with Heading 1 (0.10 points)
    try:
        toc_heading_found = False
        for h in heading1_texts:
            if 'table of contents' in h or 'toc' in h:
                toc_heading_found = True
                break
        if toc_heading_found:
            print(f"PASS: Component 2 - TOC heading found as Heading 1 (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 - No 'Table of Contents' Heading 1 found. Heading 1s: {heading1_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Six required sections with Heading 1 (0.20 points)
    # Required: Document Overview, System Description, Functional Requirements,
    #           Non-Functional Requirements, System Architecture, Glossary
    try:
        required_sections = [
            'document overview',
            'system description',
            'functional requirements',
            'non-functional requirements',
            'system architecture',
            'glossary'
        ]
        found_sections = []
        for req in required_sections:
            for h in heading1_texts:
                if req in h:
                    found_sections.append(req)
                    break

        section_ratio = len(found_sections) / len(required_sections)
        section_score = round(0.20 * section_ratio, 4)
        if section_ratio == 1.0:
            print(f"PASS: Component 3 - All 6 required sections found as Heading 1 ({section_score} pts)")
        else:
            missing = [s for s in required_sections if s not in found_sections]
            print(f"PARTIAL: Component 3 - {len(found_sections)}/6 sections found. Missing: {missing} ({section_score} pts)")
        total_score += section_score
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Functional Requirements table (0.20 points)
    # Should have 4 columns (ID, Requirement, Priority, Status) and 7 rows (1 header + 6 entries)
    try:
        fr_table = None
        for t in tables:
            # Check if first row has expected headers
            if len(t.columns) >= 4 and len(t.rows) >= 2:
                header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
                if 'id' in header_cells and 'requirement' in header_cells:
                    fr_table = t
                    break

        if fr_table is not None:
            col_count = len(fr_table.columns)
            row_count = len(fr_table.rows)
            fr_score = 0.0

            # Sub-check: correct number of columns (4)
            if col_count == 4:
                fr_score += 0.05
                print(f"  PASS: FR table has 4 columns")
            else:
                print(f"  FAIL: FR table has {col_count} columns, expected 4")

            # Sub-check: correct headers
            headers = [c.text.strip().lower() for c in fr_table.rows[0].cells]
            expected_headers = ['id', 'requirement', 'priority', 'status']
            if all(eh in headers for eh in expected_headers):
                fr_score += 0.05
                print(f"  PASS: FR table has correct headers: {headers}")
            else:
                print(f"  FAIL: FR table headers {headers} don't match expected {expected_headers}")

            # Sub-check: 7 rows (1 header + 6 data entries)
            if row_count == 7:
                fr_score += 0.05
                print(f"  PASS: FR table has 7 rows (1 header + 6 entries)")
            elif row_count >= 4:
                fr_score += 0.025
                print(f"  PARTIAL: FR table has {row_count} rows, expected 7")
            else:
                print(f"  FAIL: FR table has {row_count} rows, expected 7")

            # Sub-check: data rows have content in all cells
            data_rows_ok = 0
            for ri in range(1, min(row_count, 7)):
                cells = [c.text.strip() for c in fr_table.rows[ri].cells]
                if all(len(c) > 0 for c in cells):
                    data_rows_ok += 1
            if data_rows_ok >= 6:
                fr_score += 0.05
                print(f"  PASS: All 6 data rows have content in all cells")
            elif data_rows_ok >= 3:
                fr_score += 0.025
                print(f"  PARTIAL: {data_rows_ok}/6 data rows have content in all cells")
            else:
                print(f"  FAIL: Only {data_rows_ok} data rows have complete content")

            print(f"PASS: Component 4 - Functional Requirements table found ({fr_score} pts)")
            total_score += fr_score
        else:
            print(f"FAIL: Component 4 - No Functional Requirements table found with ID/Requirement headers")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Non-Functional Requirements - 4 bulleted items (0.15 points)
    try:
        # Find bullets after the "Non-Functional Requirements" heading
        in_nfr_section = False
        bullet_items = []
        for p in paragraphs:
            style_name = p.style.name if p.style else ''
            text = p.text.strip()

            if style_name == 'Heading 1' and 'non-functional' in text.lower():
                in_nfr_section = True
                continue
            elif style_name == 'Heading 1' and in_nfr_section:
                # Next heading => end of NFR section
                break

            if in_nfr_section and 'bullet' in style_name.lower() and len(text) > 5:
                bullet_items.append(text[:50])

        if len(bullet_items) >= 4:
            print(f"PASS: Component 5 - Found {len(bullet_items)} bulleted NFR items (0.15 pts)")
            total_score += 0.15
        elif len(bullet_items) >= 2:
            partial = round(0.15 * len(bullet_items) / 4, 4)
            print(f"PARTIAL: Component 5 - Found {len(bullet_items)}/4 bulleted NFR items ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 - Found {len(bullet_items)} bulleted items in NFR section, expected 4")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Glossary table - 2 cols, 6 rows (header + 5 terms) (0.15 points)
    try:
        glossary_table = None
        for t in tables:
            if len(t.columns) >= 2 and len(t.rows) >= 2:
                header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
                if 'term' in header_cells and 'definition' in header_cells:
                    glossary_table = t
                    break

        if glossary_table is not None:
            g_rows = len(glossary_table.rows)
            g_cols = len(glossary_table.columns)
            g_score = 0.0

            # Sub-check: correct columns
            if g_cols == 2:
                g_score += 0.05
                print(f"  PASS: Glossary table has 2 columns")
            else:
                print(f"  FAIL: Glossary table has {g_cols} columns, expected 2")

            # Sub-check: correct row count (header + 5 terms = 6 rows)
            if g_rows == 6:
                g_score += 0.05
                print(f"  PASS: Glossary table has 6 rows (header + 5 terms)")
            elif g_rows >= 4:
                g_score += 0.025
                print(f"  PARTIAL: Glossary table has {g_rows} rows, expected 6")
            else:
                print(f"  FAIL: Glossary table has {g_rows} rows, expected 6")

            # Sub-check: all terms have definitions
            terms_ok = 0
            for ri in range(1, min(g_rows, 6)):
                term = glossary_table.rows[ri].cells[0].text.strip()
                defn = glossary_table.rows[ri].cells[1].text.strip()
                if len(term) > 0 and len(defn) > 0:
                    terms_ok += 1
            if terms_ok >= 5:
                g_score += 0.05
                print(f"  PASS: All 5 glossary terms have definitions")
            elif terms_ok >= 3:
                g_score += 0.025
                print(f"  PARTIAL: {terms_ok}/5 glossary terms have definitions")
            else:
                print(f"  FAIL: Only {terms_ok} glossary terms have definitions")

            print(f"PASS: Component 6 - Glossary table found ({g_score} pts)")
            total_score += g_score
        else:
            print(f"FAIL: Component 6 - No Glossary table found with Term/Definition headers")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    # Component 7: Document Overview & System Description have body text (0.10 points)
    try:
        sections_with_body = 0
        check_sections = ['document overview', 'system description']

        for target in check_sections:
            in_section = False
            has_body = False
            for p in paragraphs:
                style_name = p.style.name if p.style else ''
                text = p.text.strip()

                if style_name == 'Heading 1' and target in text.lower():
                    in_section = True
                    continue
                elif style_name == 'Heading 1' and in_section:
                    break

                if in_section and style_name == 'Normal' and len(text) > 20:
                    has_body = True

            if has_body:
                sections_with_body += 1

        if sections_with_body == 2:
            print(f"PASS: Component 7 - Both Document Overview and System Description have body text (0.10 pts)")
            total_score += 0.10
        elif sections_with_body == 1:
            print(f"PARTIAL: Component 7 - Only 1/2 sections have body text (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 7 - Neither Document Overview nor System Description have body text")
    except Exception as e:
        print(f"ERROR: Component 7 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: persist app state then verify
def persist_app_state(domain):
    import os, time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
