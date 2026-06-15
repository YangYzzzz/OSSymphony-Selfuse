"""
Initial Setup: Trivia cards - 6 questions as plain paragraphs
Task ID: writer_creative_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_052'
OUTPUT = f'{WORKDIR}/trivia_cards.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.left_margin = Pt(72)   # 1 inch
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    # 6 trivia questions as plain paragraphs (no bold, no table, no special formatting)
    questions = [
        "Q1: What year was the first iPhone released?",
        "Q2: Which planet is known as the Red Planet?",
        "Q3: What is the capital of Australia?",
        "Q4: Who painted the Mona Lisa?",
        "Q5: What is the chemical symbol for gold?",
        "Q6: In what year did World War II end?",
    ]

    for question in questions:
        para = doc.add_paragraph()
        run = para.add_run(question)
        run.font.size = Pt(12)
        # No bold, no centering — plain left-aligned 12pt paragraphs
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
