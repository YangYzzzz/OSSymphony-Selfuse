"""
Reward Script: Layout trivia questions as a 2x3 table with bold labels and borders
Task ID: writer_creative_052
Domain: libreoffice_writer
Scoring:
  Component 1: Table exists with 2 columns and 3 rows (0.35 pts)
  Component 2: All 6 questions present in correct order in cells (0.25 pts)
  Component 3: Question labels bold at 14pt, question text at 12pt (0.25 pts)
  Component 4: Table cell borders visible (0.15 pts)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_052'

# Expected question content (label part, question part)
EXPECTED_QUESTIONS = [
    ('Q1: ', 'What year was the first iPhone released?'),
    ('Q2: ', 'Which planet is known as the Red Planet?'),
    ('Q3: ', 'What is the capital of Australia?'),
    ('Q4: ', 'Who painted the Mona Lisa?'),
    ('Q5: ', 'What is the chemical symbol for gold?'),
    ('Q6: ', 'In what year did World War II end?'),
]

def has_visible_border(cell):
    """Check if a cell has visible borders via tcBorders XML or via Table Grid style."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            # Check that at least top, left, bottom, right have a non-nil/non-none val
            for side in ['top', 'left', 'bottom', 'right']:
                el = tcBorders.find(qn(f'w:{side}'))
                if el is not None:
                    val = el.get(qn('w:val'), '')
                    if val and val not in ('nil', 'none'):
                        return True
    return False

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Table exists with exactly 2 columns and 3 rows (0.35 pts)
    # The initial state has no table (plain paragraphs), so this is the core change.
    # -----------------------------------------------------------------------
    try:
        tables = doc.tables
        if len(tables) == 0:
            print("FAIL: Component 1 — No table found in document (initial plain-paragraph state?)")
        else:
            table = tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 3 and num_cols == 2:
                print(f"PASS: Component 1 — Table has 3 rows x 2 cols (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 1 — Table is {num_rows} rows x {num_cols} cols, expected 3x2")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # For subsequent checks we need the table
    if len(doc.tables) == 0:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    table = doc.tables[0]

    # -----------------------------------------------------------------------
    # Component 2: All 6 questions present in correct cells, row-major order (0.25 pts)
    # Maps questions to cells: Q1->[0,0], Q2->[0,1], Q3->[1,0], Q4->[1,1], Q5->[2,0], Q6->[2,1]
    # -----------------------------------------------------------------------
    try:
        cells_in_order = []
        for row in table.rows:
            for cell in row.cells:
                cells_in_order.append(cell.text.strip())

        q_idx = 0
        matched = 0
        for r_idx in range(min(len(table.rows), 3)):
            for c_idx in range(min(len(table.columns), 2)):
                if q_idx >= len(EXPECTED_QUESTIONS):
                    break
                label, question_text = EXPECTED_QUESTIONS[q_idx]
                expected_full = label + question_text
                cell_text = table.cell(r_idx, c_idx).text.strip()
                if cell_text == expected_full:
                    matched += 1
                elif question_text.lower() in cell_text.lower():
                    # partial: question text present even if label differs
                    matched += 0.5
                q_idx += 1

        fraction = matched / len(EXPECTED_QUESTIONS)
        if fraction >= 1.0:
            print(f"PASS: Component 2 — All 6 questions present and correctly placed (0.25 pts)")
            total_score += 0.25
        elif fraction >= 0.5:
            partial = round(0.25 * fraction, 4)
            print(f"PARTIAL: Component 2 — {matched}/{len(EXPECTED_QUESTIONS)} questions matched, awarding {partial} pts")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched}/{len(EXPECTED_QUESTIONS)} questions matched")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Question labels bold at 14pt and question text at 12pt (0.25 pts)
    # Checks across all 6 cells. The label run should be bold=True and 14pt;
    # the question-text run should be bold=False and 12pt.
    # -----------------------------------------------------------------------
    try:
        label_ok = 0
        text_ok = 0
        total_cells = 0

        for r_idx in range(min(len(table.rows), 3)):
            for c_idx in range(min(len(table.columns), 2)):
                cell = table.cell(r_idx, c_idx)
                total_cells += 1
                for para in cell.paragraphs:
                    runs = para.runs
                    if len(runs) >= 2:
                        label_run = runs[0]
                        text_run = runs[1]
                        # Label run: bold + 14pt
                        l_bold = label_run.bold or label_run.font.bold
                        l_size = label_run.font.size.pt if label_run.font.size else None
                        if l_bold and l_size == 14.0:
                            label_ok += 1
                        # Text run: not bold, 12pt
                        t_bold = text_run.bold or text_run.font.bold
                        t_size = text_run.font.size.pt if text_run.font.size else None
                        if not t_bold and t_size == 12.0:
                            text_ok += 1
                    elif len(runs) == 1:
                        # Possibly all in one run; check if bold
                        run = runs[0]
                        run_bold = run.bold or run.font.bold
                        if run_bold:
                            label_ok += 1  # count partial credit

        if total_cells == 0:
            print("FAIL: Component 3 — No cells found to check")
        else:
            # Both conditions must hold for most cells
            if label_ok >= total_cells and text_ok >= total_cells:
                print(f"PASS: Component 3 — Label bold/14pt and text 12pt in all {total_cells} cells (0.25 pts)")
                total_score += 0.25
            elif label_ok > 0 or text_ok > 0:
                partial = round(0.25 * ((label_ok + text_ok) / (2 * total_cells)), 4)
                print(f"PARTIAL: Component 3 — label_ok={label_ok}/{total_cells}, text_ok={text_ok}/{total_cells}, awarding {partial} pts")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — No cells have correct bold/size formatting")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Table cell borders visible (0.15 pts)
    # Table Grid style or explicit tcBorders with non-nil val indicates borders.
    # -----------------------------------------------------------------------
    try:
        table_style = table.style.name if table.style else ''
        style_has_grid = ('Grid' in table_style or 'grid' in table_style)
        cell_00 = table.cell(0, 0)
        explicit_border = has_visible_border(cell_00)

        if style_has_grid:
            print(f"PASS: Component 4 — Table uses '{table_style}' style (borders implied) (0.15 pts)")
            total_score += 0.15
        elif explicit_border:
            print(f"PASS: Component 4 — Explicit cell borders found in cell [0,0] (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 — No visible borders detected (style='{table_style}', no tcBorders)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/trivia_cards.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
