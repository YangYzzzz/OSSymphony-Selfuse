"""
Initial Setup: Create a 20-page technical white paper with unformatted abstract
Task ID: writer_gf4_049
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_gf4_049'
OUTPUT = f'{WORKDIR}/white_paper.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Scalable Architectures for Distributed Machine Learning Systems: A Comprehensive Analysis', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Authors line
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Dr. Elena Vasquez, Prof. James Whitfield, Dr. Mei-Lin Chang')
    run.font.size = Pt(11)
    run.font.italic = True

    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affiliation.add_run('Institute for Advanced Computing Research, Stanford University')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Abstract (starts with "The emergence of...") ---
    abstract_heading = doc.add_heading('Abstract', level=1)

    abstract_text = (
        "The emergence of large-scale distributed computing frameworks has fundamentally "
        "transformed how organizations approach machine learning workloads. As datasets grow "
        "beyond the capacity of single-node systems, the need for efficient parallelization "
        "strategies becomes paramount. This paper presents a comprehensive analysis of modern "
        "distributed machine learning architectures, examining both data-parallel and model-parallel "
        "approaches across heterogeneous computing environments. We evaluate the performance "
        "characteristics of seven leading frameworks, including parameter server architectures, "
        "all-reduce communication patterns, and hybrid decentralized topologies. Our benchmarks "
        "span training workloads from 10 million to 175 billion parameters, demonstrating that "
        "communication overhead becomes the dominant bottleneck beyond 32-node clusters. We "
        "propose a novel adaptive gradient compression algorithm that reduces inter-node bandwidth "
        "requirements by 47% while maintaining convergence guarantees within 2% of uncompressed "
        "baselines. Furthermore, we introduce a fault-tolerant checkpoint mechanism that achieves "
        "99.7% recovery success rates with only 3.2% overhead in wall-clock training time. Our "
        "findings suggest that the optimal architecture choice depends critically on the ratio of "
        "computation to communication intensity, with model-parallel approaches favoring transformer "
        "architectures and data-parallel strategies excelling for convolutional workloads."
    )
    abstract_para = doc.add_paragraph(abstract_text)
    abstract_para.paragraph_format.space_after = Pt(12)

    # --- Keywords ---
    keywords = doc.add_paragraph()
    kr = keywords.add_run('Keywords: ')
    kr.bold = True
    kr.font.size = Pt(10)
    keywords.add_run('distributed computing, machine learning, gradient compression, '
                     'fault tolerance, scalable training, parameter servers').font.size = Pt(10)

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)

    intro_paragraphs = [
        "The rapid advancement of artificial intelligence and machine learning has created "
        "unprecedented demand for computational resources. Modern neural network architectures, "
        "particularly large language models and vision transformers, require training across "
        "thousands of GPUs for weeks or even months. This computational reality necessitates "
        "sophisticated distributed training frameworks that can efficiently utilize cluster-scale "
        "hardware while maintaining training stability and convergence properties.",

        "Historical approaches to distributed machine learning relied primarily on MapReduce-style "
        "data parallelism, where training data is partitioned across workers that independently "
        "compute gradient updates. While conceptually simple, this approach faces fundamental "
        "scaling limitations due to the synchronization overhead of aggregating gradients across "
        "an increasing number of nodes. The seminal work of Li et al. (2014) on parameter servers "
        "partially addressed this through asynchronous update mechanisms, though at the cost of "
        "potential gradient staleness and convergence instability.",

        "More recently, the emergence of dedicated interconnect technologies such as NVLink, "
        "NVSwitch, and InfiniBand has shifted the bottleneck landscape. Modern GPU clusters can "
        "achieve inter-node bandwidths exceeding 400 Gbps, enabling communication patterns that "
        "were previously infeasible. Ring all-reduce algorithms, popularized by frameworks like "
        "Horovod, exploit this bandwidth efficiently by overlapping communication with computation "
        "in carefully orchestrated pipelines.",

        "Despite these advances, significant challenges remain. The heterogeneity of modern "
        "computing environments, where nodes may contain different GPU generations or mixed "
        "CPU-GPU configurations, complicates load balancing and communication scheduling. "
        "Additionally, the increasing size of model architectures has driven interest in "
        "model-parallel strategies, where the model itself is partitioned across devices, "
        "introducing new complexities around pipeline scheduling and activation memory management.",
    ]
    for text in intro_paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Section 2: Related Work ---
    doc.add_heading('2. Related Work', level=1)

    doc.add_heading('2.1 Data Parallelism Frameworks', level=2)
    related_work_texts = [
        "Data parallelism remains the most widely adopted strategy for distributed training. "
        "The fundamental approach involves replicating the model across all workers and partitioning "
        "the training dataset. Each worker computes gradients on its local data partition, and these "
        "gradients are aggregated before applying updates to maintain model consistency. TensorFlow "
        "Distributed (Abadi et al., 2016) was among the first production-grade frameworks to support "
        "this paradigm, offering both synchronous and asynchronous training modes.",

        "PyTorch Distributed Data Parallel (DDP), introduced in version 1.0, implements bucket-based "
        "gradient all-reduce that overlaps backward computation with communication. The framework "
        "automatically groups small tensors into communication buckets, reducing the number of "
        "collective operations required per training step. Benchmarks by Li et al. (2020) demonstrate "
        "near-linear scaling up to 256 GPUs for ResNet-50 training on ImageNet.",

        "DeepSpeed, developed by Microsoft Research, extends data parallelism with ZeRO (Zero "
        "Redundancy Optimizer) stages that progressively partition optimizer states, gradients, "
        "and parameters across workers. ZeRO Stage 3 achieves memory reductions proportional to "
        "the number of workers, enabling training of models with over 100 billion parameters on "
        "commodity GPU clusters without requiring model parallelism.",
    ]
    for text in related_work_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('2.2 Model Parallelism Approaches', level=2)
    model_parallel_texts = [
        "Model parallelism addresses scenarios where a single model exceeds the memory capacity "
        "of an individual accelerator. Tensor parallelism, as implemented in Megatron-LM (Shoeybi "
        "et al., 2019), partitions individual layers across devices by splitting weight matrices "
        "along specific dimensions. For transformer architectures, the attention heads and MLP "
        "layers can be naturally partitioned, with all-reduce operations required only at layer "
        "boundaries.",

        "Pipeline parallelism takes an orthogonal approach by assigning different layers to "
        "different devices and streaming micro-batches through the pipeline. GPipe (Huang et al., "
        "2019) pioneered this approach with synchronous pipeline stages and gradient accumulation. "
        "PipeDream (Narayanan et al., 2019) extended this with asynchronous pipelines and weight "
        "stashing to maintain consistent forward and backward passes across different pipeline stages.",

        "Recent hybrid approaches combine data, tensor, and pipeline parallelism in a three-dimensional "
        "parallelism strategy. Megatron-Turing NLG 530B demonstrated that this combination can "
        "scale to thousands of GPUs while maintaining 52% of peak hardware utilization, a significant "
        "improvement over single-dimension parallelism strategies.",
    ]
    for text in model_parallel_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Section 3: Methodology ---
    doc.add_heading('3. Methodology', level=1)

    doc.add_heading('3.1 Experimental Setup', level=2)
    methodology_texts = [
        "Our experimental evaluation was conducted on a dedicated compute cluster consisting of "
        "128 nodes, each equipped with 8 NVIDIA A100-80GB GPUs interconnected via NVLink 3.0. "
        "Inter-node communication utilized a dedicated InfiniBand HDR network providing 200 Gbps "
        "per-port bandwidth. The cluster ran Ubuntu 22.04 LTS with CUDA 12.1, cuDNN 8.9, and "
        "NCCL 2.18 for collective communication primitives.",

        "We evaluated seven distributed training configurations across four model architectures: "
        "ResNet-152 (60M parameters), BERT-Large (340M parameters), GPT-2 XL (1.5B parameters), "
        "and a custom transformer variant with 13B parameters. For each configuration, we measured "
        "throughput (samples/second), time-to-accuracy (hours to reach target validation metric), "
        "communication volume (GB transferred per training step), and GPU memory utilization.",

        "To ensure reproducibility, all experiments used fixed random seeds and identical "
        "hyperparameter configurations. Each experiment was repeated three times, and we report "
        "both mean performance and standard deviation. Data loading pipelines were pre-cached to "
        "eliminate I/O variance as a confounding factor.",
    ]
    for text in methodology_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading('3.2 Gradient Compression Algorithm', level=2)
    compression_texts = [
        "We propose Adaptive Spectral Compression (ASC), a novel gradient compression algorithm "
        "that dynamically adjusts compression ratios based on the spectral properties of gradient "
        "tensors. Unlike fixed top-k or random-k sparsification, ASC performs a fast approximate "
        "SVD decomposition of gradient matrices and retains only the components whose singular "
        "values exceed a dynamically computed threshold.",

        "The threshold adaptation mechanism operates on a sliding window of gradient statistics, "
        "maintaining exponentially weighted moving averages of gradient magnitude distributions "
        "across layers. During early training phases, when gradient magnitudes are large and "
        "volatile, the algorithm retains more components to preserve training stability. As "
        "training progresses and gradients become smaller, compression ratios increase, reducing "
        "communication overhead during the computationally intensive fine-tuning phase.",

        "Error feedback mechanisms compensate for information lost during compression. We maintain "
        "a local error accumulator at each worker that adds the compression residual to the next "
        "iteration's gradient before compression. This ensures that no gradient information is "
        "permanently discarded, preserving convergence guarantees under mild assumptions on the "
        "loss landscape smoothness.",
    ]
    for text in compression_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Section 4: Results ---
    doc.add_heading('4. Results and Analysis', level=1)

    doc.add_heading('4.1 Scaling Efficiency', level=2)
    results_texts = [
        "Table 1 presents the scaling efficiency results across all framework-model combinations. "
        "For ResNet-152, all frameworks achieved near-linear scaling up to 64 GPUs, with efficiency "
        "dropping below 90% only at 128 GPUs for parameter server configurations. Ring all-reduce "
        "maintained 94.3% efficiency at 128 GPUs, while our ASC-enhanced variant achieved 96.1% "
        "by reducing communication volume by 47%.",

        "The scaling characteristics diverged significantly for larger models. GPT-2 XL training "
        "with pure data parallelism showed a marked efficiency drop beyond 32 GPUs, reaching only "
        "71.2% at 128 GPUs with standard all-reduce. The communication-intensive nature of "
        "transformer architectures, with their large embedding layers and attention matrices, "
        "makes them particularly sensitive to inter-node bandwidth limitations.",

        "Hybrid parallelism configurations demonstrated superior scaling for the 13B parameter "
        "model. A combination of 4-way tensor parallelism, 8-way data parallelism, and 4-way "
        "pipeline parallelism achieved 82.7% scaling efficiency at 128 GPUs, compared to 63.4% "
        "for pure data parallelism with ZeRO Stage 3. The reduced communication volume from "
        "tensor parallelism's intra-node all-reduce operations was the primary contributing factor.",
    ]
    for text in results_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # Add a table for results
    doc.add_heading('Table 1: Scaling Efficiency (%)', level=3)
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    headers = ['Framework', '8 GPUs', '32 GPUs', '64 GPUs', '128 GPUs']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    table_data = [
        ['PyTorch DDP', '99.1', '96.8', '93.2', '88.4'],
        ['Horovod Ring', '99.3', '97.4', '95.1', '94.3'],
        ['DeepSpeed ZeRO-3', '98.7', '95.2', '91.8', '86.9'],
        ['Megatron-LM Hybrid', '99.5', '98.1', '96.3', '82.7'],
        ['ASC (Ours)', '99.4', '97.9', '96.8', '96.1'],
    ]
    for r, row_data in enumerate(table_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')  # spacer

    doc.add_heading('4.2 Communication Analysis', level=2)
    comm_texts = [
        "Figure 2 illustrates the communication volume breakdown across training phases. "
        "During the first 10% of training steps, gradient magnitudes were 3.7x larger than "
        "the training average, resulting in higher communication volume for compressed methods. "
        "Our ASC algorithm automatically detected this regime and reduced compression ratios, "
        "maintaining training stability while still achieving a 23% bandwidth reduction compared "
        "to uncompressed baselines during this phase.",

        "The steady-state compression ratio stabilized at approximately 8.3x for ResNet-152 "
        "and 5.7x for GPT-2 XL. The difference reflects the lower rank structure of convolutional "
        "gradients compared to attention layer gradients. Importantly, ASC achieved these "
        "compression ratios without the accuracy degradation observed with fixed-ratio methods "
        "like top-k at equivalent compression levels.",

        "Network congestion analysis revealed that communication hotspots clustered around the "
        "all-reduce operations for the largest gradient tensors. By staggering compression "
        "decisions across layers and introducing a priority queue for gradient transmission, "
        "we reduced peak network utilization by 31%, improving tail latency by 2.4x at the "
        "99th percentile.",
    ]
    for text in comm_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Section 5: Fault Tolerance ---
    doc.add_heading('5. Fault Tolerance Mechanisms', level=1)
    fault_texts = [
        "Large-scale distributed training runs spanning days or weeks are susceptible to "
        "hardware failures, network partitions, and software errors. Our fault tolerance "
        "framework implements a multi-layered approach combining proactive health monitoring, "
        "elastic scaling, and efficient checkpoint recovery.",

        "The checkpoint mechanism employs an asynchronous, incremental approach. Rather than "
        "saving complete model snapshots at fixed intervals, we maintain a delta-encoded "
        "checkpoint stream that captures only the parameter changes since the last checkpoint. "
        "This reduces checkpoint size by an average of 73% and eliminates the training pause "
        "associated with synchronous checkpointing.",

        "Recovery from node failures proceeds through a three-phase protocol. First, the "
        "failure detection service identifies the failed node within 2.3 seconds on average "
        "using heartbeat monitoring with adaptive timeouts. Second, the remaining nodes "
        "redistribute the failed node's data partition and adjust gradient aggregation groups. "
        "Third, the recovered or replacement node loads the latest checkpoint and synchronizes "
        "with the current training state through a fast catch-up protocol.",

        "Our experiments demonstrated 99.7% successful recovery across 1,247 simulated failure "
        "events during 30-day continuous training runs. The average recovery time was 14.3 seconds, "
        "with 95% of recoveries completing within 28 seconds. The overhead of the fault tolerance "
        "mechanisms was measured at 3.2% of total wall-clock training time.",
    ]
    for text in fault_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Section 6: Discussion ---
    doc.add_heading('6. Discussion', level=1)
    discussion_texts = [
        "Our findings reveal several important insights for practitioners designing distributed "
        "training systems. First, the choice of parallelism strategy should be driven primarily "
        "by the computation-to-communication ratio of the target workload. Models with high "
        "arithmetic intensity, such as large convolutional networks, benefit most from data "
        "parallelism with gradient compression, while attention-heavy architectures like "
        "transformers see greater gains from tensor or pipeline parallelism.",

        "Second, the diminishing returns of scaling beyond 64 GPUs for most workloads suggest "
        "that investment in faster interconnects may yield better cost-efficiency than adding "
        "more compute nodes. Our analysis shows that doubling inter-node bandwidth from 200 Gbps "
        "to 400 Gbps would improve 128-GPU scaling efficiency by an estimated 8-12% for "
        "transformer workloads, a significant improvement relative to the infrastructure cost.",

        "Third, the interaction between gradient compression and learning rate schedules requires "
        "careful tuning. We observed that aggressive compression during warmup phases could "
        "destabilize training, while conservative compression during late training phases "
        "wasted bandwidth. The adaptive nature of ASC addresses both failure modes, but "
        "practitioners using fixed-ratio methods should implement phase-aware compression "
        "schedules.",

        "The fault tolerance results have important implications for the total cost of ownership "
        "of training clusters. With our checkpoint mechanism, the expected training time overhead "
        "from failures drops from 15-20% (with hourly synchronous checkpoints) to 3.2%, "
        "representing substantial cost savings for multi-week training runs. This overhead "
        "reduction becomes increasingly significant as cluster sizes grow, since failure "
        "probability increases super-linearly with node count.",
    ]
    for text in discussion_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Section 7: Future Work ---
    doc.add_heading('7. Future Directions', level=1)
    future_texts = [
        "Several promising directions emerge from this work. The integration of hardware-aware "
        "compression algorithms that exploit the specific communication characteristics of "
        "emerging interconnect technologies, such as CXL (Compute Express Link), could further "
        "reduce overhead. Additionally, the application of reinforcement learning to dynamically "
        "optimize parallelism configurations during training represents an exciting frontier.",

        "The emergence of sparse mixture-of-experts architectures introduces new challenges for "
        "distributed training. Unlike dense models, MoE architectures require dynamic routing "
        "of tokens to experts that may reside on different devices, creating irregular "
        "communication patterns that are poorly served by static all-reduce operations. "
        "Developing efficient communication primitives for these workloads is a critical "
        "open problem.",

        "Finally, the growing interest in federated learning and privacy-preserving distributed "
        "training opens opportunities to extend our compression and fault tolerance techniques "
        "to cross-organizational training scenarios, where communication constraints are even "
        "more severe and node reliability is less controllable.",
    ]
    for text in future_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Section 8: Conclusion ---
    doc.add_heading('8. Conclusion', level=1)
    conclusion_texts = [
        "This paper has presented a comprehensive analysis of distributed machine learning "
        "architectures, evaluating scaling efficiency, communication patterns, and fault "
        "tolerance across seven framework configurations and four model architectures. Our "
        "proposed Adaptive Spectral Compression algorithm demonstrates that intelligent gradient "
        "compression can significantly reduce communication overhead without sacrificing model "
        "quality.",

        "The key contributions of this work are threefold: (1) a systematic benchmarking "
        "methodology that enables fair comparison across heterogeneous distributed training "
        "frameworks; (2) the ASC algorithm, which achieves 47% bandwidth reduction with less "
        "than 2% accuracy impact; and (3) a fault-tolerant checkpoint mechanism that reduces "
        "training overhead from hardware failures to 3.2% of wall-clock time.",

        "As model sizes continue to grow and training compute requirements expand, the efficient "
        "utilization of distributed computing resources becomes increasingly critical. We believe "
        "that the techniques and insights presented in this work provide a foundation for the "
        "next generation of scalable training systems that can support models with trillions of "
        "parameters across thousands of heterogeneous accelerators.",
    ]
    for text in conclusion_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- References ---
    doc.add_heading('References', level=1)
    references = [
        "Abadi, M., et al. (2016). TensorFlow: A System for Large-Scale Machine Learning. OSDI.",
        "Dean, J., et al. (2012). Large Scale Distributed Deep Networks. NeurIPS.",
        "Goyal, P., et al. (2017). Accurate, Large Minibatch SGD. arXiv:1706.02677.",
        "Huang, Y., et al. (2019). GPipe: Efficient Training of Giant Neural Networks. NeurIPS.",
        "Li, M., et al. (2014). Scaling Distributed Machine Learning with the Parameter Server. OSDI.",
        "Li, S., et al. (2020). PyTorch Distributed: Experiences on Accelerating Data Parallel Training. VLDB.",
        "Narayanan, D., et al. (2019). PipeDream: Generalized Pipeline Parallelism for DNN Training. SOSP.",
        "Narayanan, D., et al. (2021). Efficient Large-Scale Language Model Training on GPU Clusters. SC.",
        "Rajbhandari, S., et al. (2020). ZeRO: Memory Optimizations Toward Training Trillion Parameter Models. SC.",
        "Shoeybi, M., et al. (2019). Megatron-LM: Training Multi-Billion Parameter Language Models. arXiv:1909.08053.",
        "Smith, S., et al. (2022). Using DeepSpeed and Megatron to Train Megatron-Turing NLG 530B. arXiv:2201.11990.",
        "Stich, S. U. (2019). Local SGD Converges Fast and Communicates Little. ICLR.",
        "You, Y., et al. (2020). Large Batch Optimization for Deep Learning: Training BERT in 76 Minutes. ICLR.",
        "Zhang, H., et al. (2017). Poseidon: An Efficient Communication Architecture for Distributed Deep Learning. ATC.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # --- Appendix A ---
    doc.add_page_break()
    doc.add_heading('Appendix A: Detailed Benchmark Configurations', level=1)
    appendix_a_texts = [
        "This appendix provides the complete hyperparameter configurations used in our "
        "experimental evaluation. All models were trained using mixed-precision (FP16) with "
        "dynamic loss scaling. Gradient clipping was set to 1.0 for all configurations.",

        "ResNet-152 Configuration: Batch size per GPU = 32, Global batch size = 256-4096 "
        "(depending on GPU count), Learning rate = 0.1 with linear warmup over 5 epochs and "
        "cosine annealing, Weight decay = 1e-4, Momentum = 0.9, Training epochs = 90, "
        "Dataset = ImageNet-1K (1.28M training images).",

        "BERT-Large Configuration: Batch size per GPU = 8, Sequence length = 512, Learning "
        "rate = 1e-4 with polynomial decay, Warmup steps = 10000, Weight decay = 0.01, "
        "Adam beta1 = 0.9, beta2 = 0.999, Training steps = 1M, Dataset = Wikipedia + BookCorpus.",

        "GPT-2 XL Configuration: Batch size per GPU = 2, Sequence length = 1024, Learning "
        "rate = 1.5e-4 with cosine decay, Warmup steps = 2000, Weight decay = 0.1, Adam "
        "beta1 = 0.9, beta2 = 0.95, Training tokens = 300B, Dataset = The Pile.",

        "Custom 13B Transformer: Batch size per GPU = 1, Sequence length = 2048, Hidden "
        "size = 5120, Attention heads = 40, Layers = 40, Learning rate = 1.2e-4 with "
        "cosine decay, Warmup steps = 3000, Weight decay = 0.1, Training tokens = 500B.",
    ]
    for text in appendix_a_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Appendix B ---
    doc.add_heading('Appendix B: Network Topology Analysis', level=1)
    appendix_b_texts = [
        "The cluster network topology consisted of a two-level fat-tree architecture with "
        "32 leaf switches, each connecting 4 nodes. The 8 spine switches provided full "
        "bisection bandwidth between any pair of leaf switches. Each node's 8 GPUs were "
        "connected via NVLink 3.0 in a fully connected topology, providing 600 GB/s aggregate "
        "intra-node GPU bandwidth.",

        "We measured actual achieved bandwidth under various communication patterns. "
        "Point-to-point transfers between GPUs on the same node achieved 95.2% of theoretical "
        "NVLink bandwidth (285 GB/s per link). Inter-node transfers via InfiniBand achieved "
        "89.7% of theoretical bandwidth (179.4 Gbps effective vs 200 Gbps theoretical).",

        "For ring all-reduce operations, the effective per-GPU bandwidth scaled as expected "
        "with the formula: B_eff = B_link * (N-1)/N, where N is the ring size. At 128 GPUs, "
        "this theoretical efficiency is 99.2%, but practical efficiency was 87.3% due to "
        "network contention and protocol overhead at the InfiniBand layer.",

        "The all-to-all communication pattern required for mixture-of-experts routing showed "
        "significantly worse scaling, achieving only 62.4% of theoretical bandwidth at 128 GPUs. "
        "This confirms our assertion that MoE architectures require fundamentally different "
        "communication optimization strategies compared to dense models.",
    ]
    for text in appendix_b_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Appendix C ---
    doc.add_heading('Appendix C: Convergence Proofs', level=1)
    appendix_c_texts = [
        "Theorem 1 (Convergence of ASC): Under standard assumptions of L-smoothness and "
        "bounded gradient variance, the ASC algorithm converges at a rate of O(1/sqrt(KT)) "
        "where K is the number of workers and T is the number of iterations, matching the "
        "convergence rate of uncompressed SGD up to constant factors.",

        "Proof Sketch: The key insight is that the spectral thresholding operation preserves "
        "the direction of the gradient update while reducing its magnitude. Combined with the "
        "error feedback mechanism, the accumulated compressed gradients asymptotically track "
        "the true gradient sum. The formal proof proceeds by bounding the variance introduced "
        "by compression using properties of the SVD decomposition and the adaptive threshold.",

        "Lemma 1: For any gradient matrix G with rank r, the ASC compression with threshold "
        "tau retains at least (1 - epsilon) of the Frobenius norm, where epsilon = "
        "sum_{i>k} sigma_i^2 / ||G||_F^2 and k is the number of retained singular values.",

        "Lemma 2: The error feedback accumulator E_t satisfies ||E_t||_F <= C * tau_t for "
        "all t, where C depends only on the loss landscape smoothness constant L and the "
        "learning rate eta. This bound ensures that the accumulated error does not grow "
        "unboundedly during training.",

        "The complete proof, including all intermediate lemmas and the extension to the "
        "non-convex setting, is available in the supplementary materials. The key technical "
        "contribution is the handling of the adaptive threshold, which introduces time-varying "
        "compression ratios that require careful analysis of the accumulator dynamics.",
    ]
    for text in appendix_c_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # --- Appendix D: Additional Results ---
    doc.add_heading('Appendix D: Additional Experimental Results', level=1)
    appendix_d_texts = [
        "This appendix contains supplementary experimental results that complement the main "
        "findings presented in Section 4. We include ablation studies, sensitivity analyses, "
        "and comparison with additional baseline methods.",

        "Ablation Study - Error Feedback: Removing the error feedback mechanism from ASC "
        "resulted in a 4.7% accuracy degradation for BERT-Large and 6.2% for GPT-2 XL at "
        "8x compression. This confirms that error feedback is essential for maintaining "
        "convergence at high compression ratios. For ResNet-152, the impact was smaller (1.3%), "
        "likely due to the higher redundancy in convolutional gradients.",

        "Ablation Study - Adaptive Threshold: Replacing the adaptive threshold with a fixed "
        "threshold calibrated to achieve the same average compression ratio resulted in training "
        "instability during the first 5% of steps. Three out of five runs with fixed thresholds "
        "diverged during warmup, compared to zero divergences with the adaptive mechanism.",

        "Sensitivity to Learning Rate: ASC showed robust performance across a 4x range of "
        "learning rates (0.5x to 2x of the optimal). At learning rates above 2x optimal, "
        "both compressed and uncompressed training diverged. At learning rates below 0.5x "
        "optimal, the convergence slowdown was proportional for both methods.",

        "Comparison with Top-K Sparsification: At equivalent compression ratios, ASC "
        "outperformed top-k sparsification by 1.2-3.8% in final accuracy across all models. "
        "The advantage was most pronounced for transformer architectures, where the low-rank "
        "structure of attention gradients aligns well with the spectral compression approach.",

        "Comparison with PowerSGD: PowerSGD achieved similar compression ratios to ASC for "
        "convolutional models but showed 1.5% accuracy degradation for transformer workloads "
        "at 128 GPUs. ASC's advantage stems from its adaptive threshold, which handles the "
        "varying gradient structure across transformer layers more effectively than PowerSGD's "
        "fixed-rank approximation.",
    ]
    for text in appendix_d_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
