"""
Reward Script: Professional column-based layout for abstract section in white_paper.docx
Task ID: writer_gf4_049
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): 2-column layout section exists around abstract
  Component 2 (0.15): Column spacing is approximately 0.5cm (283 twips)
  Component 3 (0.25): Drop cap 'T' spanning 3 lines at start of abstract
  Component 4 (0.25): Light blue (#E8F4FD) background shading on abstract paragraphs
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_gf4_049'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body

    # Gather all sectPr elements in the body (inline + final)
    all_sectPr = list(body.findall(f'.//{{{ns_w}}}sectPr'))

    # ========================================================
    # Component 1: 2-column section exists (0.35 points)
    # The task requires a 2-column layout around the abstract.
    # In OOXML, this is done via continuous section breaks with
    # cols num="2". The initial file has only 1 section, single column.
    # ========================================================
    found_2col = False
    col_space_value = None
    try:
        for sp in all_sectPr:
            cols_elem = sp.find(f'{{{ns_w}}}cols')
            if cols_elem is not None:
                num = cols_elem.get(f'{{{ns_w}}}num', '1')
                if num == '2':
                    found_2col = True
                    col_space_value = cols_elem.get(f'{{{ns_w}}}space')
                    break

        if found_2col:
            print(f"PASS: Component 1 — Found 2-column section (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — No 2-column section found. Sections: {len(all_sectPr)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ========================================================
    # Component 2: Column spacing approximately 0.5cm (0.15 points)
    # 0.5cm = 283.46 twips. Accept 270-300 range.
    # Initial file has no 2-col section so this naturally fails.
    # ========================================================
    try:
        if found_2col and col_space_value is not None:
            space_twips = int(col_space_value)
            if 270 <= space_twips <= 300:
                print(f"PASS: Component 2 — Column spacing {space_twips} twips (~0.5cm) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — Column spacing {space_twips} twips, expected ~283 (0.5cm)")
        else:
            print(f"FAIL: Component 2 — No 2-column section found, cannot check spacing")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ========================================================
    # Component 3: Drop cap 'T' spanning 3 lines (0.25 points)
    # A drop cap in OOXML is represented by a framePr element with
    # dropCap="drop" and lines="3". The text should be 'T'.
    # Initial file has no drop cap at all.
    # ========================================================
    try:
        drop_cap_found = False
        drop_cap_correct_lines = False
        drop_cap_letter = None

        for para in doc.paragraphs:
            pPr = para._element.find(f'{{{ns_w}}}pPr')
            if pPr is not None:
                framePr = pPr.find(f'{{{ns_w}}}framePr')
                if framePr is not None:
                    dc = framePr.get(f'{{{ns_w}}}dropCap')
                    if dc and dc.lower() in ('drop', 'margin'):
                        drop_cap_found = True
                        lines_attr = framePr.get(f'{{{ns_w}}}lines', '0')
                        if lines_attr == '3':
                            drop_cap_correct_lines = True
                        drop_cap_letter = para.text.strip()
                        break

        if drop_cap_found and drop_cap_correct_lines:
            # Verify the letter is 'T'
            if drop_cap_letter and drop_cap_letter.upper() == 'T':
                print(f"PASS: Component 3 — Drop cap 'T' spanning 3 lines found (0.25 pts)")
                total_score += 0.25
            else:
                # Partial: drop cap exists with correct lines but wrong letter
                print(f"FAIL: Component 3 — Drop cap found with 3 lines but letter is '{drop_cap_letter}', expected 'T'")
                total_score += 0.10
        elif drop_cap_found:
            lines_val = framePr.get(f'{{{ns_w}}}lines', 'unknown') if framePr is not None else 'unknown'
            print(f"FAIL: Component 3 — Drop cap found but lines={lines_val}, expected 3")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — No drop cap found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ========================================================
    # Component 4: Light blue (#E8F4FD) shading on abstract paragraphs (0.25 points)
    # The task requires the abstract section to have light blue background.
    # In OOXML, this is paragraph shading with fill="E8F4FD".
    # Initial file has no paragraph shading.
    # ========================================================
    try:
        shaded_paras = []
        for i, para in enumerate(doc.paragraphs):
            pPr = para._element.find(f'{{{ns_w}}}pPr')
            if pPr is not None:
                shd = pPr.find(f'{{{ns_w}}}shd')
                if shd is not None:
                    fill = shd.get(f'{{{ns_w}}}fill', '')
                    if fill and fill.lower() != 'auto':
                        shaded_paras.append((i, fill.upper(), para.text[:50]))

        if len(shaded_paras) > 0:
            # Check that at least one paragraph has the correct color
            correct_color_count = sum(1 for _, fill, _ in shaded_paras if fill == 'E8F4FD')
            if correct_color_count > 0:
                print(f"PASS: Component 4 — {correct_color_count} paragraph(s) with #E8F4FD shading (0.25 pts)")
                for idx, fill, text in shaded_paras:
                    print(f"  P{idx}: fill={fill}, text='{text}'")
                total_score += 0.25
            else:
                # Shading exists but wrong color
                print(f"FAIL: Component 4 — Shading found but with wrong color(s):")
                for idx, fill, text in shaded_paras:
                    print(f"  P{idx}: fill={fill} (expected E8F4FD)")
                total_score += 0.05
        else:
            print(f"FAIL: Component 4 — No paragraph shading found in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Main execution
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/white_paper.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
