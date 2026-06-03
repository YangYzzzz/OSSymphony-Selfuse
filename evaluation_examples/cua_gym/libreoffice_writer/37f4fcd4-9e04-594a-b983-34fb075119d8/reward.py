"""
Reward Script: Merge cells A5-C5 and type 'Section 2: Financial Analysis'
Task ID: writer_tm_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Cells A5-C5 are merged (gridSpan=3)
  Component 2 (0.35): Merged cell contains 'Section 2: Financial Analysis'
  Component 3 (0.15): Text is bold
  Component 4 (0.10): Text is centered
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_030'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as e:
        print(f"CRITICAL: Missing python-docx library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table with >= 5 rows
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) < 5:
        print(f"CRITICAL: Table has only {len(table.rows)} rows, need at least 5")
        print("REWARD: 0.0")
        return 0.0

    # Row 5 = index 4 (0-based)
    row = table.rows[4]

    # Component 1: Cells A5-C5 are merged (gridSpan=3) — 0.4 points
    try:
        cell_0 = row.cells[0]
        tc = cell_0._tc
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        gridSpan_elem = tc.find(f'{{{ns}}}tcPr/{{{ns}}}gridSpan')
        if gridSpan_elem is not None:
            span_val = int(gridSpan_elem.get(f'{{{ns}}}val', '1'))
        else:
            span_val = 1

        if span_val >= 3:
            print(f"PASS: Component 1 — Cells merged with gridSpan={span_val} (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected gridSpan>=3, found {span_val}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Merged cell contains 'Section 2: Financial Analysis' — 0.35 points
    try:
        cell_text = row.cells[0].text.strip()
        expected_text = 'Section 2: Financial Analysis'
        if expected_text.lower() in cell_text.lower():
            print(f"PASS: Component 2 — Cell text is '{cell_text}' (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 2 — Expected '{expected_text}', found '{cell_text}'")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Text is bold — 0.15 points
    try:
        cell = row.cells[0]
        has_bold_target_text = False
        for para in cell.paragraphs:
            for run in para.runs:
                if run.text.strip() and 'Section 2' in run.text:
                    if run.font.bold is True:
                        has_bold_target_text = True
                        break
            if has_bold_target_text:
                break

        if has_bold_target_text:
            print(f"PASS: Component 3 — Text is bold (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — Text is not bold")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Text is centered — 0.10 points
    try:
        cell = row.cells[0]
        is_centered = False
        for para in cell.paragraphs:
            if para.text.strip():
                alignment = para.paragraph_format.alignment
                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    is_centered = True
                    break

        if is_centered:
            print(f"PASS: Component 4 — Text is centered (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Text is not centered")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
