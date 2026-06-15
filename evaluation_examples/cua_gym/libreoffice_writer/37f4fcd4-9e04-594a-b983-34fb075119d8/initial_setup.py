"""
Initial Setup: Create a quarterly report document with a 3x12 table.
Row 5 is empty (awaiting agent to merge and add section divider).
Task ID: writer_tm_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Quarterly Business Report - Q1 2025", level=1)

    # Intro paragraph
    doc.add_paragraph(
        "This report summarizes the financial performance and operational metrics "
        "for the first quarter of 2025. The data below is organized by department "
        "with section dividers separating major categories."
    )

    # Create a 3-column x 12-row table
    table = doc.add_table(rows=12, cols=3)
    table.style = "Table Grid"

    # Row 0: Header
    headers = ["Category", "Q1 Amount ($)", "Status"]
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Section 1 header in row 1
    for col_idx in range(3):
        cell = table.cell(1, col_idx)
        if col_idx == 0:
            run = cell.paragraphs[0].add_run("Section 1: Revenue Streams")
            run.bold = True
        # Leave B1, C1 empty (unmerged section header style)

    # Rows 2-4: Revenue data
    revenue_data = [
        ["Product Sales", "245,800", "On Track"],
        ["Service Contracts", "128,450", "Above Target"],
        ["Licensing Fees", "67,320", "Below Target"],
    ]
    for r_offset, row_data in enumerate(revenue_data):
        for col_idx, val in enumerate(row_data):
            table.cell(2 + r_offset, col_idx).text = val

    # Row 5 (index 4): EMPTY - this is where the agent will merge and type
    # Leave all three cells empty
    for col_idx in range(3):
        table.cell(4, col_idx).text = ""

    # Rows 6-8: Financial analysis data
    finance_data = [
        ["Operating Expenses", "189,230", "Under Budget"],
        ["Capital Expenditures", "54,670", "On Budget"],
        ["Net Profit Margin", "12.4%", "Improving"],
    ]
    for r_offset, row_data in enumerate(finance_data):
        for col_idx, val in enumerate(row_data):
            table.cell(5 + r_offset, col_idx).text = val

    # Rows 9-11: More data
    ops_data = [
        ["Employee Headcount", "342", "Stable"],
        ["Customer Acquisition", "1,287", "Growing"],
        ["Support Tickets Resolved", "4,563", "Improved"],
    ]
    for r_offset, row_data in enumerate(ops_data):
        for col_idx, val in enumerate(row_data):
            table.cell(8 + r_offset, col_idx).text = val

    # Row 12 (index 11): Summary
    summary_data = ["Total Revenue", "441,570", "On Track"]
    for col_idx, val in enumerate(summary_data):
        cell = table.cell(11, col_idx)
        run = cell.paragraphs[0].add_run(val)
        run.bold = True

    # Closing paragraph
    doc.add_paragraph("")
    doc.add_paragraph(
        "For detailed breakdowns, please refer to the attached appendices."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
