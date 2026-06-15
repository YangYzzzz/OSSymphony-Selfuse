"""
Reward Script: Meeting minutes template in LibreOffice Writer
Task ID: writer_wf_005
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Info table — 2 columns, 5 rows with correct field labels
  Component 2 (0.3): Four Heading 2 sections (Agenda Items, Discussion Points, Action Items, Next Meeting)
  Component 3 (0.2): Action Items table — 3 columns with headers Action, Responsible, Deadline
  Component 4 (0.2): Liberation Sans 11pt font applied consistently
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_005'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Info table — 2 columns, 5 rows with correct field labels (0.3 points)
    # The task requires a table at the top with 2 columns and 5 rows containing specific fields.
    try:
        expected_labels = ['Meeting Title:', 'Date:', 'Time:', 'Location:', 'Attendees:']
        if len(doc.tables) >= 1:
            info_table = doc.tables[0]
            rows_ok = len(info_table.rows) == 5
            cols_ok = len(info_table.columns) == 2

            # Check that column 0 has the correct labels
            found_labels = []
            for ri, row in enumerate(info_table.rows):
                cell_text = row.cells[0].text.strip()
                found_labels.append(cell_text)

            labels_match = found_labels == expected_labels

            if rows_ok and cols_ok and labels_match:
                print(f"PASS: Component 1 — Info table 2x5 with correct labels (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — rows_ok={rows_ok}, cols_ok={cols_ok}, labels_match={labels_match}")
                print(f"  Found labels: {found_labels}")
                # Partial credit: table exists with approximately right structure
                if rows_ok and cols_ok:
                    partial = 0.0
                    for i, label in enumerate(expected_labels):
                        if i < len(found_labels) and label.lower() in found_labels[i].lower():
                            partial += 0.06
                    if partial > 0:
                        print(f"  Partial credit: {partial:.2f} pts for partially matching labels")
                        total_score += partial
        else:
            print(f"FAIL: Component 1 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Four Heading 2 sections (0.3 points)
    # The task requires sections with Heading 2: 'Agenda Items', 'Discussion Points',
    # 'Action Items', 'Next Meeting'
    try:
        expected_headings = ['Agenda Items', 'Discussion Points', 'Action Items', 'Next Meeting']
        found_headings = []
        for para in doc.paragraphs:
            if para.style and para.style.name == 'Heading 2':
                found_headings.append(para.text.strip())

        matching = 0
        for eh in expected_headings:
            if any(eh.lower() == fh.lower() for fh in found_headings):
                matching += 1

        if matching == 4:
            print(f"PASS: Component 2 — All 4 Heading 2 sections found (0.3 pts)")
            total_score += 0.3
        elif matching > 0:
            partial = round(matching * 0.075, 3)
            print(f"FAIL: Component 2 — Found {matching}/4 headings. Partial credit: {partial} pts")
            print(f"  Found headings: {found_headings}")
            if partial > 0:
                total_score += partial
        else:
            print(f"FAIL: Component 2 — No Heading 2 sections found. Found headings: {found_headings}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Action Items table with 3 columns: Action, Responsible, Deadline (0.2 points)
    # This should be a separate table (not the info table), with the correct column headers.
    try:
        comp3_score = 0.0
        if len(doc.tables) >= 2:
            # Check tables beyond the first for the action items table
            for t_idx in range(1, len(doc.tables)):
                tbl = doc.tables[t_idx]
                if len(tbl.columns) == 3 and len(tbl.rows) >= 1:
                    header_cells = [tbl.rows[0].cells[c].text.strip() for c in range(3)]
                    expected_headers = ['Action', 'Responsible', 'Deadline']
                    if all(eh.lower() == hc.lower() for eh, hc in zip(expected_headers, header_cells)):
                        comp3_score = 0.2
                        break

            if comp3_score > 0:
                print(f"PASS: Component 3 — Action Items table with correct 3 columns (0.2 pts)")
                total_score += comp3_score
            else:
                # Check if any table has 3 columns with partially matching headers
                for t_idx in range(1, len(doc.tables)):
                    tbl = doc.tables[t_idx]
                    if len(tbl.columns) == 3:
                        header_cells = [tbl.rows[0].cells[c].text.strip() for c in range(3)]
                        print(f"FAIL: Component 3 — Found 3-col table but headers are {header_cells}")
                        break
                else:
                    print(f"FAIL: Component 3 — No 3-column table found beyond the info table")
        else:
            print(f"FAIL: Component 3 — Only {len(doc.tables)} table(s), need at least 2")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Liberation Sans 11pt font throughout (0.2 points)
    # Check that runs with explicit font settings use Liberation Sans at 11pt.
    try:
        total_runs_with_font = 0
        matching_font_runs = 0

        # Check paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name is not None:
                    total_runs_with_font += 1
                    name_ok = 'liberation sans' in run.font.name.lower()
                    size_ok = run.font.size is not None and abs(run.font.size.pt - 11.0) < 0.5
                    if name_ok and size_ok:
                        matching_font_runs += 1

        # Check table cells
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.name is not None:
                                total_runs_with_font += 1
                                name_ok = 'liberation sans' in run.font.name.lower()
                                size_ok = run.font.size is not None and abs(run.font.size.pt - 11.0) < 0.5
                                if name_ok and size_ok:
                                    matching_font_runs += 1

        if total_runs_with_font == 0:
            print(f"FAIL: Component 4 — No runs with explicit font settings found")
        elif matching_font_runs == total_runs_with_font:
            print(f"PASS: Component 4 — All {total_runs_with_font} styled runs use Liberation Sans 11pt (0.2 pts)")
            total_score += 0.2
        elif matching_font_runs > 0:
            ratio = matching_font_runs / total_runs_with_font
            partial = round(0.2 * ratio, 3)
            print(f"FAIL: Component 4 — {matching_font_runs}/{total_runs_with_font} runs match. Partial: {partial} pts")
            if partial > 0:
                total_score += partial
        else:
            print(f"FAIL: Component 4 — 0/{total_runs_with_font} runs use Liberation Sans 11pt")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice edits before verification
def persist_app_state(domain):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
