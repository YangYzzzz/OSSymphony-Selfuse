"""
Reward Script: Insert bibliography entry for RFC 7231 and add bibliography table
Task ID: writer_tech_047
Domain: libreoffice_writer
Scoring:
  Component 1: Citation mark [RFC7231] in body text (0.25 pts)
  Component 2: Bibliography heading at end of document (0.15 pts)
  Component 3: Bibliography table exists with correct structure (0.30 pts)
  Component 4: Bibliography table content correctness (0.30 pts)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_047'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Citation mark [RFC7231] present in body text (0.25 points)
    # In the initial document, the text says "RFC 7231 - HTTP/1.1 Semantics and Content."
    # with NO [RFC7231] citation mark. The golden adds "[RFC7231]" as a citation.
    try:
        citation_found = False
        for i, para in enumerate(doc.paragraphs):
            # Skip the bibliography section itself (last few paragraphs)
            # We only want citations in the body text
            if para.style.name in ('Heading 1',) and para.text.strip().lower() in ('bibliography', 'references'):
                break
            if '[RFC7231]' in para.text or '[RFC 7231]' in para.text:
                citation_found = True
                print(f"PASS: Component 1 - Citation mark found in paragraph {i}: '...{para.text[-80:]}'")
                break

        # Also check for CITATION field code in the XML as an alternative form
        if not citation_found:
            from lxml import etree
            for elem in doc.element.body.iter():
                if elem.tag.endswith('instrText') and elem.text and 'CITATION' in elem.text and 'RFC7231' in elem.text:
                    citation_found = True
                    print(f"PASS: Component 1 - CITATION field code found: {repr(elem.text)}")
                    break

        if citation_found:
            total_score += 0.25
        else:
            print("FAIL: Component 1 - No [RFC7231] citation mark found in body text")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Bibliography heading exists at end of document (0.15 points)
    # Initial doc has "References" heading; golden has "Bibliography" heading.
    # We accept either "Bibliography" as the heading, since that's the task requirement.
    try:
        bib_heading_found = False
        # Look for a heading containing "Bibliography" in the last portion of the document
        for i in range(len(doc.paragraphs) - 1, max(len(doc.paragraphs) - 15, -1), -1):
            para = doc.paragraphs[i]
            if para.style.name.startswith('Heading') and 'bibliograph' in para.text.lower():
                bib_heading_found = True
                print(f"PASS: Component 2 - Bibliography heading found at P{i}: '{para.text}' (style: {para.style.name})")
                break

        if bib_heading_found:
            total_score += 0.15
        else:
            print("FAIL: Component 2 - No 'Bibliography' heading found near end of document")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Bibliography table exists with correct structure (0.30 points)
    # Initial doc has 1 table (HTTP methods). Golden has 2 tables (methods + bibliography).
    # The bibliography table should have 2 columns (Field, Value) and multiple rows
    # with fields like Identifier, Type, Author, Title, Publisher, Year, etc.
    try:
        bib_table = None
        # Look for a table that has bibliography-like structure (2 cols, field/value pairs)
        for table in doc.tables:
            if len(table.columns) == 2 and len(table.rows) >= 5:
                # Check if this looks like a bibliography table
                cell_texts = [row.cells[0].text.strip().lower() for row in table.rows]
                if any(kw in cell_texts for kw in ['identifier', 'author', 'title', 'year', 'type']):
                    bib_table = table
                    break

        if bib_table is not None:
            num_rows = len(bib_table.rows)
            print(f"PASS: Component 3 - Bibliography table found with {num_rows} rows and 2 columns")
            # Give partial credit based on how many expected fields are present
            expected_fields = {'identifier', 'author', 'title', 'year', 'publisher', 'type'}
            found_fields = set()
            for row in bib_table.rows:
                field_name = row.cells[0].text.strip().lower()
                if field_name in expected_fields:
                    found_fields.add(field_name)
            field_ratio = len(found_fields) / len(expected_fields)
            points = 0.30 * field_ratio
            if field_ratio < 1.0:
                print(f"  Partial: Found fields {found_fields}, missing {expected_fields - found_fields} ({field_ratio:.0%})")
            total_score += points
            print(f"  Score: {points:.2f}/0.30 pts")
        else:
            print("FAIL: Component 3 - No bibliography table found with expected structure")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Bibliography table content correctness (0.30 points)
    # Check that key bibliography values for RFC 7231 are correct
    try:
        if bib_table is not None:
            # Build a dict of field -> value from the table
            bib_data = {}
            for row in bib_table.rows:
                field = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                bib_data[field] = value

            points_4 = 0.0
            checks = 0
            max_checks = 5

            # Check 1: Identifier contains RFC7231 or RFC 7231
            if 'identifier' in bib_data:
                if 'rfc7231' in bib_data['identifier'].lower().replace(' ', '').replace('-', ''):
                    points_4 += 0.30 / max_checks
                    checks += 1
                    print(f"  PASS: Identifier = '{bib_data['identifier']}'")
                else:
                    print(f"  FAIL: Identifier = '{bib_data['identifier']}', expected RFC7231")
            else:
                print("  FAIL: No 'Identifier' field in bibliography table")

            # Check 2: Author contains Fielding
            if 'author' in bib_data:
                if 'fielding' in bib_data['author'].lower():
                    points_4 += 0.30 / max_checks
                    checks += 1
                    print(f"  PASS: Author = '{bib_data['author']}'")
                else:
                    print(f"  FAIL: Author = '{bib_data['author']}', expected Fielding")
            else:
                print("  FAIL: No 'Author' field in bibliography table")

            # Check 3: Title contains HTTP/1.1 and Semantics
            if 'title' in bib_data:
                title_lower = bib_data['title'].lower()
                if 'http' in title_lower and 'semantic' in title_lower:
                    points_4 += 0.30 / max_checks
                    checks += 1
                    print(f"  PASS: Title = '{bib_data['title']}'")
                else:
                    print(f"  FAIL: Title = '{bib_data['title']}', expected HTTP/1.1 Semantics")
            else:
                print("  FAIL: No 'Title' field in bibliography table")

            # Check 4: Year is 2014
            if 'year' in bib_data:
                if '2014' in bib_data['year']:
                    points_4 += 0.30 / max_checks
                    checks += 1
                    print(f"  PASS: Year = '{bib_data['year']}'")
                else:
                    print(f"  FAIL: Year = '{bib_data['year']}', expected 2014")
            else:
                print("  FAIL: No 'Year' field in bibliography table")

            # Check 5: Publisher contains IETF
            if 'publisher' in bib_data:
                if 'ietf' in bib_data['publisher'].lower():
                    points_4 += 0.30 / max_checks
                    checks += 1
                    print(f"  PASS: Publisher = '{bib_data['publisher']}'")
                else:
                    print(f"  FAIL: Publisher = '{bib_data['publisher']}', expected IETF")
            else:
                print("  FAIL: No 'Publisher' field in bibliography table")

            total_score += points_4
            print(f"PASS: Component 4 - {checks}/{max_checks} content checks passed ({points_4:.2f}/0.30 pts)")
        else:
            print("FAIL: Component 4 - No bibliography table to check content")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
