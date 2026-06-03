"""
FINAL REWARD SCRIPT - SUCCESS
Task: Paragraph 9 is full of hard-tab characters and the spacing looks off. In LibreOffice Writer, how can I isolate that single paragraph and swap every “	” with exactly two regular space characters (‘  ’) in one go?
Generated: 2025-09-10 14:55:59
Status: success
Model: azure-o3
Total Steps: 5
"""

import os
import re
from docx import Document

def verify_task(file_path: str) -> float:
    """Verify that paragraph 9 no longer contains hard-tabs (\t) and that
    each former tab has been replaced with EXACTLY two regular spaces.

    Scoring (progressive – totals to 1.0):
        • 0.4 – Paragraph 9 contains **no** tab characters.
        • 0.4 – Every inter-word space sequence in paragraph 9 is **exactly** two spaces.
        • 0.2 – No tab characters remain **anywhere** in the document.
    """

    max_score = 1.0
    score = 0.0

    # ---------- prerequisite: file must exist and be loadable ----------
    if not os.path.exists(file_path):
        print("✗ File not found:", file_path)
        return 0.0

    try:
        doc = Document(file_path)
        print(f"✓ Loaded DOCX – {len(doc.paragraphs)} paragraphs found")
    except Exception as e:
        print(f"✗ Could not load DOCX: {e}")
        return 0.0

    # ---------- ensure document has at least 9 paragraphs ----------
    if len(doc.paragraphs) < 9:
        print("✗ Document has fewer than 9 paragraphs – cannot verify paragraph 9")
        return 0.0

    paragraph9 = doc.paragraphs[8].text  # 0-based index
    print("Paragraph 9 text:")
    print(repr(paragraph9))

    # ---------- Requirement 1: no hard-tab characters in paragraph 9 ----------
    if "\t" not in paragraph9:
        print("✓ Paragraph 9 has no hard-tab characters (0.4)")
        score += 0.4
    else:
        print("✗ Paragraph 9 still contains hard-tabs")

    # ---------- Requirement 2: each tab replaced by EXACTLY two spaces ----------
    # Capture every run of spaces occurring between two non-space characters
    space_sequences = [m.group(2) for m in re.finditer(r"(\S+)( +)(?=\S)", paragraph9)]

    if space_sequences:
        lengths = [len(s) for s in space_sequences]
        unique_lengths = set(lengths)
        print("Inter-word space lengths detected:", lengths)
        if unique_lengths == {2}:
            print("✓ All inter-word space sequences are exactly 2 spaces (0.4)")
            score += 0.4
        else:
            print(f"✗ Space sequences are not uniform 2-spaces – lengths found: {sorted(unique_lengths)}")
    else:
        # If there is only one word (unlikely) we cannot prove replacement was done.
        print("✗ No inter-word space sequences detected – cannot verify replacement")

    # ---------- Requirement 3: no hard-tabs anywhere in the document ----------
    tabs_elsewhere = [idx + 1 for idx, p in enumerate(doc.paragraphs) if "\t" in p.text]

    if not tabs_elsewhere:
        print("✓ No hard-tab characters found in any paragraph (0.2)")
        score += 0.2
    else:
        print("✗ Hard-tabs still present in paragraph(s):", tabs_elsewhere)

    final_score = min(score, max_score)
    print(f"Total Score: {final_score}/{max_score}")
    return final_score

# -------------------- EXECUTION ENTRY POINT --------------------
if __name__ == "__main__":
    FILE_PATH = "/home/user/paragraph_9_is_full_of_hard_tab_characters_and_the_spacing_looks_off_in_libreoffice_writer_how_can_i.docx"
    reward = verify_task(FILE_PATH)
    print(f"REWARD: {reward}")
