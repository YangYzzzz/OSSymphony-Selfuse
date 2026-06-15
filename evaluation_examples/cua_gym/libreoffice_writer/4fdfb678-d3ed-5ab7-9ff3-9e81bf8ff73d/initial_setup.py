"""
Initial Setup: Insert a Table of Figures in a document with 8 captioned figures
Task ID: writer_mt_055
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Figure caption texts (the descriptive part after "Figure N: ")
FIGURE_CAPTIONS = [
    "Experimental Setup",
    "Sample Preparation",
    "Temperature Calibration",
    "Pressure Readings",
    "Chemical Analysis",
    "Microscope Imaging",
    "Data Distribution",
    "Final Results",
]

# Realistic section content for each figure
SECTION_CONTENT = [
    "The experimental apparatus was assembled according to the standard protocol outlined in Section 2.1 of the methodology guide. The primary reaction vessel was connected to the temperature-controlled water bath, with digital thermocouples positioned at three monitoring points along the condenser column.",
    "Samples were prepared using a double-blind randomization procedure. Each specimen was weighed to the nearest 0.001 g using an analytical balance (Mettler Toledo XPE205) and then dissolved in 50 mL of deionized water at room temperature. The resulting solutions were filtered through 0.45 micrometer membrane filters before analysis.",
    "Temperature calibration was performed using a three-point reference method with certified thermometric standards at 0.0, 25.0, and 100.0 degrees Celsius. The calibration curve showed excellent linearity with an R-squared value of 0.9998, confirming the accuracy of our measurement system.",
    "Pressure readings were recorded at 15-minute intervals over a 6-hour observation period. The initial pressure was set to 101.3 kPa and allowed to equilibrate for 30 minutes before data collection began. Fluctuations remained within the acceptable tolerance range of plus or minus 0.5 kPa throughout the experiment.",
    "Chemical analysis was conducted using high-performance liquid chromatography (HPLC) with a C18 reversed-phase column. The mobile phase consisted of a 60:40 acetonitrile-water mixture with 0.1% trifluoroacetic acid. Detection was performed at 254 nm wavelength using a UV-Vis detector.",
    "Microscope imaging was performed using a Zeiss Axio Observer Z1 inverted microscope equipped with differential interference contrast (DIC) optics. Images were captured at 40x and 100x magnification using a Hamamatsu ORCA-Flash4.0 digital camera with 2048 by 2048 pixel resolution.",
    "The data distribution analysis revealed a bimodal pattern in the measured concentrations. The primary peak centered around 45.2 mg/L with a standard deviation of 3.8 mg/L, while the secondary peak appeared at 78.6 mg/L. A Shapiro-Wilk test confirmed non-normal distribution at the 95% confidence level.",
    "Final results demonstrate that the treatment protocol achieved a 94.7% removal efficiency for the target contaminant. Statistical analysis using a paired t-test confirmed a significant difference between pre-treatment and post-treatment concentrations with a p-value less than 0.001.",
]


def add_fld_char(run, fld_type):
    """Add a w:fldChar element to a run."""
    fld = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): fld_type})
    run._element.append(fld)


def add_instr_text(run, text):
    """Add a w:instrText element to a run."""
    instr = run._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = text
    run._element.append(instr)


def add_caption_paragraph(doc, figure_num, caption_text, style_name='Caption'):
    """
    Add a properly formatted figure caption with SEQ field code.
    Creates: "Figure N: Caption Text" where N is a SEQ Figure field.
    """
    para = doc.add_paragraph()
    para.style = doc.styles['Caption']
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # "Figure " prefix
    run_prefix = para.add_run("Figure ")

    # SEQ Figure field: begin
    run_begin = para.add_run()
    add_fld_char(run_begin, 'begin')

    # SEQ Figure field: instruction
    run_instr = para.add_run()
    add_instr_text(run_instr, ' SEQ Figure \\* ARABIC ')

    # SEQ Figure field: separate
    run_sep = para.add_run()
    add_fld_char(run_sep, 'separate')

    # SEQ Figure field: cached display value
    run_val = para.add_run(str(figure_num))

    # SEQ Figure field: end
    run_end = para.add_run()
    add_fld_char(run_end, 'end')

    # ": Caption Text"
    run_caption = para.add_run(f": {caption_text}")

    return para


def create_placeholder_image(path, width=640, height=400, label=""):
    """Create a simple placeholder image for figures."""
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new('RGB', (width, height), color=(240, 240, 245))
    draw = ImageDraw.Draw(img)

    # Draw border
    draw.rectangle([2, 2, width - 3, height - 3], outline=(180, 180, 190), width=2)

    # Draw grid lines for a "chart-like" appearance
    for x in range(0, width, width // 8):
        draw.line([(x, 0), (x, height)], fill=(220, 220, 230), width=1)
    for y in range(0, height, height // 6):
        draw.line([(0, y), (width, y)], fill=(220, 220, 230), width=1)

    # Draw some random-ish data bars/points depending on figure
    import hashlib
    seed = int(hashlib.md5(label.encode()).hexdigest()[:8], 16)
    bar_count = 8
    bar_width = (width - 80) // bar_count
    colors = [(70, 130, 180), (60, 179, 113), (255, 165, 0), (220, 20, 60),
              (147, 112, 219), (0, 191, 255), (255, 215, 0), (50, 205, 50)]
    for i in range(bar_count):
        h = 40 + ((seed * (i + 1) * 7) % (height - 120))
        x0 = 40 + i * bar_width + 5
        x1 = x0 + bar_width - 10
        y0 = height - 40
        y1 = y0 - h
        color = colors[i % len(colors)]
        draw.rectangle([x0, y1, x1, y0], fill=color)

    # Add label text
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
    except (OSError, IOError):
        font = ImageFont.load_default()
    if label:
        bbox = draw.textbbox((0, 0), label, font=font)
        tw = bbox[2] - bbox[0]
        draw.text(((width - tw) // 2, 10), label, fill=(60, 60, 80), font=font)

    img.save(path)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Laboratory Analysis Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle / Author info ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Department of Analytical Chemistry\nResearch Division B — Q1 2025")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacing

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report presents the findings from the comprehensive laboratory analysis "
        "conducted between January and March 2025. The study evaluated the effectiveness "
        "of a novel contaminant removal protocol using advanced spectroscopic and "
        "chromatographic techniques. Results indicate a statistically significant "
        "improvement in removal efficiency compared to the baseline method, achieving "
        "94.7% contaminant removal under controlled conditions."
    )

    # --- Introduction ---
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Environmental contamination monitoring requires precise analytical methods "
        "capable of detecting trace-level concentrations across diverse sample matrices. "
        "The present study was initiated following preliminary field observations that "
        "suggested elevated contaminant levels in the watershed adjacent to the "
        "industrial processing facility. This investigation follows established EPA "
        "Method 524.2 protocols with modifications for enhanced sensitivity."
    )
    doc.add_paragraph(
        "The objectives of this study were threefold: (1) to characterize the baseline "
        "contamination profile, (2) to evaluate the performance of the proposed treatment "
        "methodology under laboratory conditions, and (3) to develop a predictive model "
        "for treatment optimization based on the experimental data collected."
    )

    # --- Methods and Results with Figures ---
    doc.add_heading("Methods and Results", level=1)

    for i, (caption_text, content) in enumerate(zip(FIGURE_CAPTIONS, SECTION_CONTENT)):
        fig_num = i + 1
        section_title = caption_text
        doc.add_heading(f"Section {fig_num}: {section_title}", level=2)
        doc.add_paragraph(content)

        # Create and insert placeholder image
        img_path = f'{WORKDIR}/fig_{fig_num}.png'
        create_placeholder_image(img_path, label=f"Figure {fig_num}: {caption_text}")
        doc.add_picture(img_path, width=Inches(5.0))
        # Center the image paragraph
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add caption with SEQ field
        add_caption_paragraph(doc, fig_num, caption_text)

        doc.add_paragraph()  # spacing

    # --- Conclusion ---
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The experimental results presented in this report demonstrate the viability "
        "of the proposed treatment protocol for contaminant removal. The 94.7% removal "
        "efficiency significantly exceeds the regulatory threshold of 85% established "
        "by the environmental protection guidelines. Further field-scale validation is "
        "recommended before full implementation."
    )

    # --- References ---
    doc.add_heading("References", level=1)
    references = [
        "Anderson, R.J., & Thompson, K.L. (2024). Advanced Spectroscopic Methods for Environmental Analysis. Journal of Analytical Chemistry, 156(3), 234-248.",
        "Chen, W., Martinez, P.D., & O'Brien, S.F. (2023). Treatment Optimization Using Statistical Design of Experiments. Environmental Science & Technology, 57(12), 4892-4901.",
        "EPA Method 524.2 (2022). Measurement of Purgeable Organic Compounds in Water by Capillary Column Gas Chromatography/Mass Spectrometry. U.S. Environmental Protection Agency.",
        "Harrison, G.M., & Patel, N.K. (2024). Membrane Filtration Technologies for Water Purification. Water Research, 248, 120871.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


create_initial()
