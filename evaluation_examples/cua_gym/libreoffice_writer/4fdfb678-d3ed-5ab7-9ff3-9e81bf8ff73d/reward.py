"""
Reward Script: Insert a Table of Figures listing all captioned images
Task ID: writer_mt_055
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): TOC field code with Figure category exists
  Component 2 (0.20): "Table of Figures" heading paragraph present
  Component 3 (0.30): All 8 figure captions listed in TOF entries
  Component 4 (0.20): Figure entries include page numbers
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_055'

# The 8 expected figure captions (from the task context)
EXPECTED_FIGURES = [
    "Figure 1: Experimental Setup",
    "Figure 2: Sample Preparation",
    "Figure 3: Temperature Calibration",
    "Figure 4: Pressure Readings",
    "Figure 5: Chemical Analysis",
    "Figure 6: Microscope Imaging",
    "Figure 7: Data Distribution",
    "Figure 8: Final Results",
]


def persist_app_state(domain: str):
    """Save any unsaved edits in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from lxml import etree
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # ---------------------------------------------------------------
    # Component 1: TOC field code with Figure category exists (0.30)
    # A Table of Figures in docx uses a TOC field with \c "Figure".
    # The initial doc has NO TOC field; the golden doc has one.
    # ---------------------------------------------------------------
    try:
        all_instr = doc.element.body.findall('.//w:instrText', ns)
        toc_figure_found = False
        for instr in all_instr:
            if instr.text and 'TOC' in instr.text and '"Figure"' in instr.text:
                toc_figure_found = True
                break
            # Also handle cases where \c Figure is without quotes or with single quotes
            if instr.text and 'TOC' in instr.text and '\\c' in instr.text and 'Figure' in instr.text:
                toc_figure_found = True
                break

        if toc_figure_found:
            print(f"PASS: Component 1 -- TOC field with Figure category found (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 -- No TOC field with Figure category found")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # ---------------------------------------------------------------
    # Component 2: "Table of Figures" heading paragraph present (0.20)
    # The golden doc has a Heading 1 or similar heading with text
    # "Table of Figures". The initial doc does NOT have this.
    # ---------------------------------------------------------------
    try:
        tof_heading_found = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            style_name = para.style.name if para.style else ''
            # Accept any heading style or "TOC Heading" style
            if 'table of figures' in text_lower and ('heading' in style_name.lower() or 'toc' in style_name.lower()):
                tof_heading_found = True
                break

        if tof_heading_found:
            print(f"PASS: Component 2 -- 'Table of Figures' heading found (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 -- No 'Table of Figures' heading found")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # ---------------------------------------------------------------
    # Component 3: All 8 figure captions listed in TOF entries (0.30)
    # The TOF should contain entries for all 8 figures.
    # We look for paragraphs (not Caption-styled) that contain figure
    # caption text. Each found figure earns 0.30/8 = 0.0375 points.
    # Initial doc has no such TOF paragraphs.
    # ---------------------------------------------------------------
    try:
        # Collect all non-Caption paragraph texts that contain "Figure N:"
        # Caption paragraphs are the actual captions in the body, not TOF entries
        tof_entry_texts = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            text = para.text.strip()
            # TOF entries are typically Normal or "Table of Figures" style, NOT Caption
            if style_name != 'Caption' and re.match(r'Figure\s+\d+:', text):
                tof_entry_texts.append(text)

        figures_found = 0
        for expected in EXPECTED_FIGURES:
            # Check if any TOF entry starts with the expected caption text
            matched = any(entry.startswith(expected) for entry in tof_entry_texts)
            if matched:
                figures_found += 1

        points_per_figure = 0.30 / 8.0
        component3_score = figures_found * points_per_figure

        if figures_found == 8:
            print(f"PASS: Component 3 -- All 8 figures found in Table of Figures (0.30 pts)")
        elif figures_found > 0:
            print(f"PARTIAL: Component 3 -- {figures_found}/8 figures found in TOF ({component3_score:.4f} pts)")
        else:
            print(f"FAIL: Component 3 -- No figure entries found in Table of Figures")

        total_score += component3_score
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # ---------------------------------------------------------------
    # Component 4: Figure entries include page numbers (0.20)
    # Each TOF entry should have a tab followed by a page number.
    # The format is "Figure N: Caption\tPageNum".
    # Initial doc has no such entries, so this naturally scores 0.
    # ---------------------------------------------------------------
    try:
        entries_with_page = 0
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            text = para.text.strip()
            if style_name != 'Caption' and re.match(r'Figure\s+\d+:', text):
                # Check if there's a tab followed by a number (page number)
                if re.search(r'\t\d+', text):
                    entries_with_page += 1

        points_per_entry = 0.20 / 8.0
        component4_score = entries_with_page * points_per_entry

        if entries_with_page == 8:
            print(f"PASS: Component 4 -- All 8 entries have page numbers (0.20 pts)")
        elif entries_with_page > 0:
            print(f"PARTIAL: Component 4 -- {entries_with_page}/8 entries have page numbers ({component4_score:.4f} pts)")
        else:
            print(f"FAIL: Component 4 -- No TOF entries with page numbers found")

        total_score += component4_score
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: persist state then verify
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
