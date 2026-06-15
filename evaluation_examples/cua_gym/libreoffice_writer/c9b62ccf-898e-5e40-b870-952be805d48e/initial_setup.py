"""
Initial Setup: Create HR salary table document (unsorted)
Task ID: writer_tm_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('HR Salary Report', level=1)

    # Add a brief intro paragraph
    doc.add_paragraph(
        'The following table summarizes the current salary information '
        'for employees across various departments.'
    )

    # Create 4x8 table (1 header + 7 data rows), unsorted
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'

    # Headers
    headers = ['Name', 'Title', 'Department', 'Salary']
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Data rows - intentionally NOT sorted by salary
    data = [
        ['Rachel Kim',       'Software Engineer',     'Engineering',  65000],
        ['David Okonkwo',    'Product Manager',       'Product',      82000],
        ['Maria Santos',     'HR Coordinator',        'Human Resources', 55000],
        ['James Thornton',   'VP of Sales',           'Sales',        91000],
        ['Priya Patel',      'Data Analyst',          'Analytics',    73000],
        ['Liam Chen',        'Junior Designer',       'Design',       48000],
        ['Aisha Washington', 'Senior Accountant',     'Finance',      87000],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            if col_idx == 3:  # Salary column - store as text with dollar format
                cell.text = str(val)
            else:
                cell.text = str(val)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
