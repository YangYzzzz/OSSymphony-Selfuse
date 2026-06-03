"""
Reward Script: Sort HR salary table by Salary column descending
Task ID: writer_tm_026
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Salary column values in correct descending order
  Component 2 (0.3): Name column values in correct sorted order
  Component 3 (0.2): First data row fully matches expected top earner
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_026'

# Expected order after sorting by Salary descending
EXPECTED_SALARIES = ['91000', '87000', '82000', '73000', '65000', '55000', '48000']
EXPECTED_NAMES = [
    'James Thornton',
    'Aisha Washington',
    'David Okonkwo',
    'Priya Patel',
    'Rachel Kim',
    'Maria Santos',
    'Liam Chen',
]
EXPECTED_FIRST_ROW = ['James Thornton', 'VP of Sales', 'Sales', '91000']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table with 8 rows and 4 columns
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) < 8 or len(table.columns) < 4:
        print(f"CRITICAL: Table dimensions wrong — {len(table.rows)} rows x {len(table.columns)} cols, expected 8x4")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: header row must still be intact
    headers = [table.cell(0, c).text.strip() for c in range(4)]
    if headers != ['Name', 'Title', 'Department', 'Salary']:
        print(f"CRITICAL: Header row corrupted — found {headers}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Salary column values in correct descending order (0.5 points)
    try:
        actual_salaries = [table.cell(r, 3).text.strip() for r in range(1, 8)]
        if actual_salaries == EXPECTED_SALARIES:
            print(f"PASS: Component 1 — Salary column in correct descending order {actual_salaries} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Expected salaries {EXPECTED_SALARIES}, found {actual_salaries}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Name column values in correct sorted order (0.3 points)
    try:
        actual_names = [table.cell(r, 0).text.strip() for r in range(1, 8)]
        if actual_names == EXPECTED_NAMES:
            print(f"PASS: Component 2 — Name column in correct sorted order (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected names {EXPECTED_NAMES}, found {actual_names}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: First data row fully matches top earner (0.2 points)
    try:
        actual_first_row = [table.cell(1, c).text.strip() for c in range(4)]
        if actual_first_row == EXPECTED_FIRST_ROW:
            print(f"PASS: Component 3 — First data row matches top earner {actual_first_row} (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Expected first data row {EXPECTED_FIRST_ROW}, found {actual_first_row}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
