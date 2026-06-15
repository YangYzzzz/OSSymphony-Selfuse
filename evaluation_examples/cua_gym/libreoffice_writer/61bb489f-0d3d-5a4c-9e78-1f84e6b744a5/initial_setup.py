"""
Initial Setup: Create a business report with a References section (no hanging indents)
Task ID: writer_biz_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document styles setup ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('Quarterly Market Analysis Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by the Strategic Research Division')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('March 2026')
    run.font.size = Pt(11)
    run.font.italic = True

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents findings from our comprehensive analysis of market trends '
        'across the Asia-Pacific technology sector during Q1 2026. Key indicators suggest '
        'a sustained growth trajectory in cloud computing, artificial intelligence, and '
        'enterprise automation. Our research methodology combined quantitative market data '
        'with qualitative assessments from industry stakeholders.'
    )
    doc.add_paragraph(
        'Revenue projections for the sector indicate a 14.3% year-over-year increase, '
        'driven primarily by accelerated digital transformation initiatives among mid-market '
        'enterprises. The competitive landscape continues to evolve, with emerging players '
        'capturing market share in specialized AI services.'
    )

    # --- Market Overview ---
    doc.add_heading('Market Overview', level=1)
    doc.add_paragraph(
        'The Asia-Pacific technology market reached a valuation of $847 billion in Q1 2026, '
        'representing a significant expansion from $741 billion in the same period last year. '
        'Cloud infrastructure spending increased by 22%, while enterprise software licensing '
        'grew at a more moderate 9.7%. Mobile commerce platforms showed particular strength '
        'in Southeast Asian markets.'
    )
    doc.add_paragraph(
        'Notable trends include the rapid adoption of generative AI tools across financial '
        'services, healthcare, and manufacturing sectors. Regulatory frameworks in several '
        'countries have begun to crystallize, providing clearer guidance for AI deployment '
        'in customer-facing applications.'
    )

    # --- Key Findings ---
    doc.add_heading('Key Findings', level=1)
    doc.add_paragraph(
        'Cloud computing revenue: $312 billion (+22% YoY)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'AI/ML services market: $98.4 billion (+31% YoY)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Enterprise automation: $67.2 billion (+18% YoY)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Cybersecurity spending: $54.8 billion (+25% YoY)',
        style='List Bullet'
    )

    # --- Methodology ---
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(
        'Data was collected through a combination of primary interviews with 142 industry '
        'executives, analysis of publicly available financial reports from 85 technology '
        'companies, and proprietary survey data from 2,300 enterprise IT decision-makers '
        'across 12 countries in the Asia-Pacific region.'
    )

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The Asia-Pacific technology sector demonstrates robust growth fundamentals heading '
        'into Q2 2026. Organizations that strategically invest in AI capabilities and cloud '
        'infrastructure modernization are positioned to capture disproportionate value. We '
        'recommend continued monitoring of regulatory developments and competitive dynamics.'
    )

    # --- References section (NO hanging indent - task requires adding it) ---
    doc.add_heading('References', level=1)

    references = [
        'Anderson, T., & Nakamura, K. (2025). Cloud Computing Trends in the Asia-Pacific Region: '
        'A Comprehensive Market Analysis. Journal of Technology Management, 42(3), 215-234. '
        'https://doi.org/10.1234/jtm.2025.0042',

        'Chen, W., Patel, R., & Okonkwo, A. (2026). Enterprise AI Adoption Patterns Across '
        'Emerging Markets: Evidence from Southeast Asia. International Review of Information '
        'Technology, 18(1), 78-99. https://doi.org/10.5678/irit.2026.0018',

        'Fujimoto, H. (2025). Regulatory Frameworks for Artificial Intelligence in Financial '
        'Services: A Comparative Study of APAC Jurisdictions. Asia-Pacific Law Review, 31(4), '
        '401-428. https://doi.org/10.9012/aplr.2025.0031',

        'Gupta, S., & Lee, M. (2026). Cybersecurity Investment Strategies for Mid-Market '
        'Enterprises: Balancing Risk and Innovation. Cybersecurity Quarterly, 9(1), 33-51. '
        'https://doi.org/10.3456/csq.2026.0009',

        'Kim, J., Zhang, L., & Ramirez, D. (2025). Digital Transformation and Organizational '
        'Performance: A Meta-Analysis of 200 Case Studies. Management Science Letters, 15(6), '
        '887-912. https://doi.org/10.7890/msl.2025.0015',

        'Morrison, E., & Tanaka, Y. (2025). Mobile Commerce Platforms in Southeast Asia: Growth '
        'Drivers and Consumer Behavior Patterns. E-Commerce Research Journal, 22(2), 145-170. '
        'https://doi.org/10.2345/ecrj.2025.0022',

        'Osei, B., Singh, A., & Watanabe, T. (2026). The Impact of Generative AI on Enterprise '
        'Productivity: Quantitative Evidence from a Multi-Country Study. Artificial Intelligence '
        'Review, 55(1), 23-48. https://doi.org/10.6789/air.2026.0055',

        'Petrova, N., & Huang, X. (2025). Cloud Infrastructure Scaling Patterns and Cost '
        'Optimization in High-Growth Technology Companies. IEEE Transactions on Cloud Computing, '
        '13(4), 562-581. https://doi.org/10.1109/tcc.2025.0013',
    ]

    for ref in references:
        para = doc.add_paragraph(ref)
        # Explicitly set NO indent to ensure clean initial state
        para.paragraph_format.left_indent = Inches(0)
        para.paragraph_format.first_line_indent = Inches(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
