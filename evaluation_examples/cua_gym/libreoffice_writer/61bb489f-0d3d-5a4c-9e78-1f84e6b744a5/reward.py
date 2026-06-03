"""
Reward Script: Hanging indent for References section
Task ID: writer_biz_043
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Reference paragraphs have left_indent ~0.5 inches
  Component 2 (0.4): Reference paragraphs have first_line_indent ~-0.5 inches (hanging)
  Component 3 (0.2): Non-reference paragraphs are unaffected
"""

import os
from docx import Document
from docx.shared import Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_043'

# Tolerance for indent checks (in inches)
TOLERANCE = 0.05
EXPECTED_LEFT_INDENT = 0.5    # inches
EXPECTED_FIRST_LINE = -0.5    # inches (negative = hanging)


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def find_references_section(doc):
    """Find the index of the 'References' heading and return indices of reference paragraphs."""
    ref_heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') and para.text.strip() == 'References':
            ref_heading_idx = i
            break

    if ref_heading_idx is None:
        return None, []

    # Reference paragraphs are all paragraphs after the heading until end or next heading
    ref_para_indices = []
    for i in range(ref_heading_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style.name.startswith('Heading'):
            break
        if para.text.strip():  # non-empty paragraphs
            ref_para_indices.append(i)

    return ref_heading_idx, ref_para_indices


def emu_to_inches(emu_val):
    """Convert EMU to inches. Returns None if input is None."""
    if emu_val is None:
        return None
    return emu_val / 914400.0


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ref_heading_idx, ref_para_indices = find_references_section(doc)
    if ref_heading_idx is None:
        print("CRITICAL: No 'References' heading found in document")
        print("REWARD: 0.0")
        return 0.0

    if len(ref_para_indices) == 0:
        print("CRITICAL: No reference paragraphs found after 'References' heading")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found 'References' heading at paragraph {ref_heading_idx}")
    print(f"INFO: Found {len(ref_para_indices)} reference paragraphs: {ref_para_indices}")

    num_refs = len(ref_para_indices)

    # Component 1: Reference paragraphs have left_indent ~0.5 inches (0.4 points)
    try:
        left_indent_pass = 0
        for idx in ref_para_indices:
            para = doc.paragraphs[idx]
            li = emu_to_inches(para.paragraph_format.left_indent)
            if li is not None and abs(li - EXPECTED_LEFT_INDENT) <= TOLERANCE:
                left_indent_pass += 1
            else:
                print(f"FAIL: Component 1 — P{idx} left_indent={li}, expected ~{EXPECTED_LEFT_INDENT}")

        if left_indent_pass == num_refs:
            print(f"PASS: Component 1 — All {num_refs} reference paragraphs have left_indent ~0.5in (0.4 pts)")
            total_score += 0.4
        elif left_indent_pass > 0:
            partial = round(0.4 * left_indent_pass / num_refs, 2)
            print(f"PARTIAL: Component 1 — {left_indent_pass}/{num_refs} paragraphs correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No reference paragraphs have correct left_indent")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Reference paragraphs have first_line_indent ~-0.5 inches (0.4 points)
    try:
        fli_pass = 0
        for idx in ref_para_indices:
            para = doc.paragraphs[idx]
            fli = emu_to_inches(para.paragraph_format.first_line_indent)
            if fli is not None and abs(fli - EXPECTED_FIRST_LINE) <= TOLERANCE:
                fli_pass += 1
            else:
                print(f"FAIL: Component 2 — P{idx} first_line_indent={fli}, expected ~{EXPECTED_FIRST_LINE}")

        if fli_pass == num_refs:
            print(f"PASS: Component 2 — All {num_refs} reference paragraphs have first_line_indent ~-0.5in (0.4 pts)")
            total_score += 0.4
        elif fli_pass > 0:
            partial = round(0.4 * fli_pass / num_refs, 2)
            print(f"PARTIAL: Component 2 — {fli_pass}/{num_refs} paragraphs correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No reference paragraphs have correct first_line_indent")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Hanging indent applied ONLY to references, not the whole document (0.2 points)
    # Compound check: at least one reference paragraph has correct hanging indent AND
    # no non-reference paragraphs have been given the hanging indent.
    # This ensures the agent targeted only the References section.
    try:
        # Sub-check A: count reference paragraphs that have the hanging indent
        ref_indent_count = sum(
            1 for idx in ref_para_indices
            if (emu_to_inches(doc.paragraphs[idx].paragraph_format.left_indent) is not None
                and abs(emu_to_inches(doc.paragraphs[idx].paragraph_format.left_indent) - EXPECTED_LEFT_INDENT) <= TOLERANCE
                and emu_to_inches(doc.paragraphs[idx].paragraph_format.first_line_indent) is not None
                and abs(emu_to_inches(doc.paragraphs[idx].paragraph_format.first_line_indent) - EXPECTED_FIRST_LINE) <= TOLERANCE)
        )

        # Sub-check B: no non-reference paragraphs have the hanging indent
        non_ref_indices = [i for i in range(len(doc.paragraphs))
                          if i not in ref_para_indices and i != ref_heading_idx]
        contaminated = 0
        for idx in non_ref_indices:
            para = doc.paragraphs[idx]
            li = emu_to_inches(para.paragraph_format.left_indent)
            fli = emu_to_inches(para.paragraph_format.first_line_indent)
            if (li is not None and abs(li - EXPECTED_LEFT_INDENT) <= TOLERANCE and
                    fli is not None and abs(fli - EXPECTED_FIRST_LINE) <= TOLERANCE):
                contaminated += 1
                print(f"FAIL: Component 3 — P{idx} has hanging indent but should not (text: {para.text[:50]!r})")

        if ref_indent_count > 0 and contaminated == 0:
            print(f"PASS: Component 3 — Hanging indent applied only to reference paragraphs (0.2 pts)")
            total_score += 0.2
        elif ref_indent_count == 0:
            print(f"FAIL: Component 3 — No reference paragraphs have hanging indent yet")
        else:
            print(f"FAIL: Component 3 — {contaminated} non-reference paragraphs incorrectly have hanging indent")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
