"""
Initial Setup: Project status report with 2 headings and 3 body paragraphs, all at single spacing.
Task ID: osworld_writer_line_spacing_per_paragraph_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Heading 1: Project Overview ---
    h1 = doc.add_heading("Project Overview", level=1)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)

    # --- Body paragraph 1 ---
    body1 = doc.add_paragraph(
        "The Q1 2025 infrastructure modernization project is currently on track to meet all "
        "scheduled milestones. The team has successfully completed the initial assessment phase "
        "and is now transitioning into the core implementation stage. Key stakeholders have been "
        "briefed and have provided their approval for the proposed architecture changes."
    )
    body1.paragraph_format.line_spacing = 1.0
    body1.paragraph_format.space_before = Pt(0)
    body1.paragraph_format.space_after = Pt(0)

    # --- Body paragraph 2 ---
    body2 = doc.add_paragraph(
        "Resource allocation remains well-managed with no critical shortfalls identified at this "
        "time. The engineering team, led by Sarah Chen, has maintained consistent velocity over "
        "the past four sprints. External vendor contracts with TechPartners Inc. and CloudBase "
        "Solutions have been finalized as of February 14, 2025, ensuring timely delivery of "
        "required components."
    )
    body2.paragraph_format.line_spacing = 1.0
    body2.paragraph_format.space_before = Pt(0)
    body2.paragraph_format.space_after = Pt(0)

    # --- Heading 2: Risk Assessment ---
    h2 = doc.add_heading("Risk Assessment", level=1)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)

    # --- Body paragraph 3 ---
    body3 = doc.add_paragraph(
        "Three moderate-risk items have been identified and are being actively monitored. "
        "The potential delay in legacy system data migration has been flagged by Marcus Johnson "
        "as requiring immediate attention. Contingency plans have been drafted and reviewed "
        "by the project steering committee. The risk register has been updated as of "
        "March 3, 2025, to reflect the latest status across all workstreams."
    )
    body3.paragraph_format.line_spacing = 1.0
    body3.paragraph_format.space_before = Pt(0)
    body3.paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
