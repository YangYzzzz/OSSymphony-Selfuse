"""
Initial Setup: Create a blank Writer document with legal-themed content.
Task ID: writer_frd_055
Domain: libreoffice_writer
Initial state: LibreOffice Writer open with a document, default AutoText categories only.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    heading = doc.add_heading('Contract Standardization Initiative', level=1)

    # Introductory paragraph
    doc.add_paragraph(
        'As part of our ongoing effort to streamline the legal review process, '
        'the Legal Operations team has identified several recurring clauses that '
        'appear across virtually all of our commercial agreements. Standardizing '
        'these clauses will reduce drafting time and ensure consistency across '
        'all contracts issued by the firm.'
    )

    doc.add_heading('Identified Clauses for Standardization', level=2)

    doc.add_paragraph(
        'The following three clauses have been selected for the initial phase '
        'of the standardization project:'
    )

    # Bullet list of the clauses to standardize
    doc.add_paragraph('Force Majeure Clause', style='List Bullet')
    doc.add_paragraph('Indemnification Clause', style='List Bullet')
    doc.add_paragraph('Governing Law Clause', style='List Bullet')

    doc.add_heading('Action Items', level=2)

    doc.add_paragraph(
        'Each clause should be added as an AutoText entry in LibreOffice Writer '
        'under a new category called "Legal Templates" so that all attorneys and '
        'paralegals can quickly insert standardized language into new contracts. '
        'The shortcuts should be concise and memorable for rapid insertion.'
    )

    doc.add_paragraph(
        'Please coordinate with the Senior Partner to finalize the exact wording '
        'of each clause before adding them to the template library. Once approved, '
        'these entries should be accessible to all team members through the '
        'Tools > AutoText dialog (Ctrl+F3).'
    )

    doc.add_heading('Timeline', level=2)

    doc.add_paragraph(
        'Target completion date: April 15, 2026. All three AutoText entries '
        'must be created and verified by this date.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure the default autotext directory exists and only has defaults
    autotext_dir = '/home/user/.config/libreoffice/4/user/autotext'
    os.makedirs(autotext_dir, exist_ok=True)
    # Remove any leftover custom autotext .bau files (except mytexts.bau which is default)
    for fname in os.listdir(autotext_dir):
        if fname.endswith('.bau') and fname != 'mytexts.bau':
            os.remove(os.path.join(autotext_dir, fname))
            print(f'Removed stale autotext file: {fname}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
