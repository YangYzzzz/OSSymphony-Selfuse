"""
Reward Script: Avery 5395 Name Badge Label Layout
Task ID: writer_lec_053
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Table with 4 rows x 2 cols (8 badges)
  Component 2 (0.35): Each badge has 'VISITOR' in bold ~20pt
  Component 3 (0.20): Each badge has a name line/space below VISITOR
  Component 4 (0.15): Page margins reduced from default 1-inch
"""

import os
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_053'


def verify_task(file_path):
    """
    Verify Avery 5395 name badge label layout.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table with 4 rows x 2 cols = 8 badges (0.30 points)
    try:
        tables = doc.tables
        if len(tables) >= 1:
            table = tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            total_badges = num_rows * num_cols
            if num_rows == 4 and num_cols == 2:
                print(f"PASS: Component 1 -- Table is 4x2 = 8 badges (0.30 pts)")
                total_score += 0.30
            elif total_badges == 8:
                # Accept alternate layouts that still give 8 badges (e.g. 2x4, 8x1)
                print(f"PASS: Component 1 -- Table has {total_badges} badges ({num_rows}x{num_cols}) (0.30 pts)")
                total_score += 0.30
            elif total_badges >= 6:
                # Partial credit for close-enough badge count
                print(f"PARTIAL: Component 1 -- Table has {total_badges} badges ({num_rows}x{num_cols}), expected 8 (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 -- Table has {total_badges} badges ({num_rows}x{num_cols}), expected 8")
        else:
            print(f"FAIL: Component 1 -- No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Each badge has 'VISITOR' in bold, ~20pt font (0.35 points)
    try:
        if len(tables) >= 1:
            table = tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            total_cells = num_rows * num_cols
            visitor_bold_count = 0

            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    cell_has_visitor_bold = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text = run.text.strip().upper()
                            if 'VISITOR' in text:
                                is_bold = run.font.bold is True
                                font_size_pt = run.font.size.pt if run.font.size else None
                                # Accept font sizes in the range 16-28pt as "large"
                                size_ok = font_size_pt is not None and 16 <= font_size_pt <= 28
                                # Mark as valid if bold OR large size (style-inherited bold possible)
                                cell_has_visitor_bold = is_bold or size_ok
                    if cell_has_visitor_bold:
                        visitor_bold_count += 1

            if total_cells > 0:
                ratio = visitor_bold_count / total_cells
                if ratio >= 1.0:
                    print(f"PASS: Component 2 -- All {visitor_bold_count}/{total_cells} badges have VISITOR bold/large (0.35 pts)")
                    total_score += 0.35
                elif ratio >= 0.5:
                    pts = round(0.35 * ratio, 2)
                    print(f"PARTIAL: Component 2 -- {visitor_bold_count}/{total_cells} badges have VISITOR bold/large ({pts} pts)")
                    total_score += pts
                else:
                    print(f"FAIL: Component 2 -- Only {visitor_bold_count}/{total_cells} badges have VISITOR bold/large")
            else:
                print(f"FAIL: Component 2 -- No table cells to check")
        else:
            print(f"FAIL: Component 2 -- No tables found")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Each badge has a name line/space below VISITOR (0.20 points)
    # Check that each cell has content beyond just "VISITOR" -- e.g. underscores, "NAME", blank line
    try:
        if len(tables) >= 1:
            table = tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            total_cells = num_rows * num_cols
            name_space_count = 0

            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    # Check if cell has more than just VISITOR
                    # Look for underscores, "NAME", blank lines, or multiple paragraphs
                    has_visitor = 'VISITOR' in cell_text.upper()
                    if not has_visitor:
                        continue

                    # Remove "VISITOR" and check remaining content
                    remaining = cell_text.upper().replace('VISITOR', '').strip()
                    # Check for name space indicators: underscores, "NAME", multiple paragraphs
                    has_underscore = '_' in cell_text
                    has_name_label = 'NAME' in remaining
                    has_multiple_paras = len([p for p in cell.paragraphs if True]) >= 2

                    if has_underscore or has_name_label or has_multiple_paras:
                        name_space_count += 1

            if total_cells > 0:
                ratio = name_space_count / total_cells
                if ratio >= 1.0:
                    print(f"PASS: Component 3 -- All {name_space_count}/{total_cells} badges have name space (0.20 pts)")
                    total_score += 0.20
                elif ratio >= 0.5:
                    pts = round(0.20 * ratio, 2)
                    print(f"PARTIAL: Component 3 -- {name_space_count}/{total_cells} badges have name space ({pts} pts)")
                    total_score += pts
                else:
                    print(f"FAIL: Component 3 -- Only {name_space_count}/{total_cells} badges have name space")
            else:
                print(f"FAIL: Component 3 -- No table cells")
        else:
            print(f"FAIL: Component 3 -- No tables found")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Page margins adjusted from default 1-inch (0.15 points)
    # Initial file has 1-inch (914400 EMU) margins all around.
    # Golden has reduced margins. Check that at least one margin is smaller than default.
    try:
        section = doc.sections[0]
        default_margin = 914400  # 1 inch in EMU
        lm = section.left_margin or default_margin
        rm = section.right_margin or default_margin
        tm = section.top_margin or default_margin
        bm = section.bottom_margin or default_margin

        margins_changed = 0
        for name, val in [('left', lm), ('right', rm), ('top', tm), ('bottom', bm)]:
            if val < default_margin:
                margins_changed += 1

        if margins_changed >= 2:
            print(f"PASS: Component 4 -- {margins_changed} margins reduced from default (0.15 pts)")
            total_score += 0.15
        elif margins_changed >= 1:
            print(f"PARTIAL: Component 4 -- {margins_changed} margin(s) reduced from default (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 4 -- All margins are still at default 1 inch (L={lm} R={rm} T={tm} B={bm})")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
