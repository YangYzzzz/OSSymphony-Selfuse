"""
FINAL REWARD SCRIPT - SUCCESS
Task: Insert a 7x5 table at the cursor and keep it left-aligned with the text.
Generated: 2025-10-14 08:58:56
Status: success
Model: azure-o3
Total Steps: 2
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

"""
Reward Script for Writer Task
Task: "Insert a 7x5 table at the cursor and keep it left-aligned with the text."

Verification Logic:
1. Load the specified DOCX file (no points – prerequisite).
2. Search all tables for one that matches 7 rows x 5 columns (or the transpose 5 x 7, as some
   tools store dimensions reversed). 0.6 points.
3. For a qualifying table, verify it is left-aligned. 0.4 points.
4. Return a progressive score, capped at 1.0.

Key Anti-Bias Measures:
• No points for merely loading the file or for tables that do not meet the size requirement.
• Alignment points are only awarded if the correct-size table exists.
"""

EXPECTED_DIMENSIONS = {(7, 5), (5, 7)}  # Accept either orientation
FILE_PATH = "/home/user/insert_a_7x5_table_at_the_cursor_and_keep_it_left_aligned_with_the_text.docx"

def is_left_aligned(table):
    """Return True if table alignment is default (None) or explicitly LEFT."""
    align = table.alignment
    return align in (None, WD_TABLE_ALIGNMENT.LEFT)

def verify_table_task(file_path):
    print(f"Verifying document: {file_path}")
    score = 0.0  # progressive score

    # --- Prerequisite: load file (no points) ---
    if not os.path.exists(file_path):
        print("✗ File not found")
        return 0.0
    try:
        doc = Document(file_path)
        print(f"✓ Loaded DOCX. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
    except Exception as exc:
        print(f"✗ Failed to load DOCX: {exc}")
        return 0.0

    # --- Requirement 1: 7×5 table present ---
    correct_table_found = False
    left_aligned_found = False

    for idx, tbl in enumerate(doc.tables):
        rows, cols = len(tbl.rows), len(tbl.columns)
        print(f"  Table {idx+1}: {rows} rows × {cols} columns")

        if (rows, cols) in EXPECTED_DIMENSIONS:
            print("   ✓ Matches 7×5 dimension requirement")
            correct_table_found = True

            if is_left_aligned(tbl):
                print("   ✓ Table is left-aligned")
                left_aligned_found = True
            else:
                print("   ✗ Table is not left-aligned")

    # --- Scoring ---
    if correct_table_found:
        score += 0.6  # main size requirement fulfilled
    else:
        print("✗ No 7×5 table found")

    if correct_table_found and left_aligned_found:
        score += 0.4  # alignment requirement fulfilled
    elif correct_table_found:
        print("✗ Correct table found, but alignment incorrect")

    score = min(score, 1.0)
    print(f"Final Score: {score}")
    return score

if __name__ == "__main__":
    reward = verify_table_task(FILE_PATH)
    print(f"REWARD: {reward}")

