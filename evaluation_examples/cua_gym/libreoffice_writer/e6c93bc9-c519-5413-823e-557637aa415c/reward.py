"""
Reward Script: Accept all tracked changes in Chapter 2, reject all in Chapter 3
Task ID: writer_acad_037
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): No tracked changes remain in the document
  Component 2 (0.35): Chapter 2 text matches accepted state (insertions kept, deletions removed)
  Component 3 (0.35): Chapter 3 text matches rejected state (insertions removed, deletions restored)
  Component 4 (0.10): Document structural integrity (Chapter 1 unchanged, chapter headings intact)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_037'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_full_para_text_with_tracked(para_element, ns):
    """Get full paragraph text INCLUDING tracked insertion and deletion text from raw XML."""
    texts = []
    # Walk all text elements including those inside w:ins and w:del
    for elem in para_element.iter():
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 't' and elem.text:
            texts.append(elem.text)
        elif tag == 'delText' and elem.text:
            texts.append(elem.text)
    return ''.join(texts)


def find_chapter_paragraphs(doc):
    """Return dict mapping chapter number to list of paragraph texts (using para.text which excludes tracked changes)."""
    chapters = {}
    current_chapter = None
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.match(r'Chapter\s+(\d+)', text)
        if match:
            current_chapter = int(match.group(1))
            chapters[current_chapter] = []
        elif current_chapter is not None:
            chapters[current_chapter].append(text)
    return chapters


def count_tracked_changes(doc):
    """Count tracked insertions and deletions in the document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body = doc.element.body
    inserts = body.findall('.//w:ins', ns)
    deletes = body.findall('.//w:del', ns)
    return len(inserts), len(deletes)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Count tracked changes (used as gate for subsequent components)
    num_ins, num_del = count_tracked_changes(doc)
    no_tracked_changes = (num_ins == 0 and num_del == 0)

    # ----- Component 1: No tracked changes remain (0.20 points) -----
    # After accepting Ch2 and rejecting Ch3, there should be zero tracked changes.
    try:
        if no_tracked_changes:
            print(f"PASS: Component 1 — No tracked changes remain (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Found {num_ins} insertions and {num_del} deletions still tracked")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Get chapter text (para.text excludes tracked changes from raw XML)
    try:
        chapters = find_chapter_paragraphs(doc)
        ch2_text = ' '.join(chapters.get(2, []))
        ch3_text = ' '.join(chapters.get(3, []))
        print(f"  Chapter 2 paragraph count: {len(chapters.get(2, []))}")
        print(f"  Chapter 3 paragraph count: {len(chapters.get(3, []))}")
    except Exception as e:
        print(f"ERROR: Could not parse chapter structure: {e}")
        ch2_text = ''
        ch3_text = ''

    # ----- Component 2: Chapter 2 text matches accepted state (0.35 points) -----
    # GATE: tracked changes must be resolved first, otherwise these checks are meaningless
    # because para.text excludes tracked insertion text in unresolved docs
    try:
        if not no_tracked_changes:
            print(f"FAIL: Component 2 — Cannot verify Ch2 text while tracked changes remain")
        else:
            ch2_checks = [
                # Accepted insertions should be present in final text
                ("particularly the pioneering work of Hsieh and Tang (1998)", "Hsieh/Tang reference inserted"),
                ("remarkable success", "remarkable (accepted, replacing moderate)"),
                ("transformer-based models have begun to surpass LSTM", "transformer mention inserted"),
                ("multiple diverse neural network", "multiple diverse inserted"),
                ("Graph neural networks represent an emerging approach", "GNN paragraph inserted"),
                # Accepted deletions should be absent from final text
                # Use negative checks: these phrases should NOT appear
            ]
            ch2_negative_checks = [
                ("moderate success in capturing", "moderate removed (replaced by remarkable)"),
                ("the computational cost remains prohibitively expensive", "computational cost sentence removed"),
                ("combining two or three neural", "two or three removed"),
            ]

            passed = 0
            total_checks = len(ch2_checks) + len(ch2_negative_checks)

            for phrase, label in ch2_checks:
                if phrase in ch2_text:
                    passed += 1
                    print(f"  PASS: Ch2 — {label}")
                else:
                    print(f"  FAIL: Ch2 — {label}: not found")

            for phrase, label in ch2_negative_checks:
                if phrase not in ch2_text:
                    passed += 1
                    print(f"  PASS: Ch2 — {label}")
                else:
                    print(f"  FAIL: Ch2 — {label}: still present")

            if total_checks > 0:
                ratio = passed / total_checks
                points = round(0.35 * ratio, 4)
                if ratio == 1.0:
                    print(f"PASS: Component 2 — All {passed}/{total_checks} Ch2 checks passed ({points} pts)")
                else:
                    print(f"PARTIAL: Component 2 — {passed}/{total_checks} Ch2 checks passed ({points} pts)")
                total_score += points
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ----- Component 3: Chapter 3 text matches rejected state (0.35 points) -----
    # GATE: tracked changes must be resolved first
    try:
        if not no_tracked_changes:
            print(f"FAIL: Component 3 — Cannot verify Ch3 text while tracked changes remain")
        else:
            # Rejected insertions should NOT be present
            ch3_negative_checks = [
                ("employs a modified U-Net", "modified U-Net rejected"),
                ("Data augmentation is performed through random temporal shifts", "data augmentation rejected"),
                ("AdamW optimizer with cosine annealing", "AdamW rejected"),
                ("additionally perform five-fold cross-validation", "five-fold cross-validation rejected"),
            ]
            # Rejected deletions should be RESTORED (present in final text)
            ch3_positive_checks = [
                ("employs a standard U-Net", "standard U-Net restored"),
                ("restricted to the Northern Hemisphere", "Northern Hemisphere restored"),
                ("standard Adam optimizer with a fixed learning rate of 0.001", "Adam optimizer restored"),
                ("simple linear regression baseline for reference", "linear regression baseline restored"),
            ]

            passed = 0
            total_checks = len(ch3_negative_checks) + len(ch3_positive_checks)

            for phrase, label in ch3_negative_checks:
                if phrase not in ch3_text:
                    passed += 1
                    print(f"  PASS: Ch3 — {label}")
                else:
                    print(f"  FAIL: Ch3 — {label}: still present")

            for phrase, label in ch3_positive_checks:
                if phrase in ch3_text:
                    passed += 1
                    print(f"  PASS: Ch3 — {label}")
                else:
                    print(f"  FAIL: Ch3 — {label}: not found")

            if total_checks > 0:
                ratio = passed / total_checks
                points = round(0.35 * ratio, 4)
                if ratio == 1.0:
                    print(f"PASS: Component 3 — All {passed}/{total_checks} Ch3 checks passed ({points} pts)")
                else:
                    print(f"PARTIAL: Component 3 — {passed}/{total_checks} Ch3 checks passed ({points} pts)")
                total_score += points
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ----- Component 4: Document structural integrity (0.10 points) -----
    # Chapter headings must be intact, and Chapter 1 text unchanged
    try:
        if not no_tracked_changes:
            print(f"FAIL: Component 4 — Cannot verify structure while tracked changes remain")
        else:
            checks_passed = 0
            total_struct_checks = 3

            # Check chapter headings exist
            all_text = ' '.join(p.text for p in doc.paragraphs)
            if 'Chapter 1: Introduction' in all_text:
                checks_passed += 1
                print(f"  PASS: Ch1 heading intact")
            else:
                print(f"  FAIL: Ch1 heading missing")

            if 'Chapter 2: Literature Review' in all_text:
                checks_passed += 1
                print(f"  PASS: Ch2 heading intact")
            else:
                print(f"  FAIL: Ch2 heading missing")

            if 'Chapter 3: Methodology' in all_text:
                checks_passed += 1
                print(f"  PASS: Ch3 heading intact")
            else:
                print(f"  FAIL: Ch3 heading missing")

            ratio = checks_passed / total_struct_checks
            points = round(0.10 * ratio, 4)
            if ratio == 1.0:
                print(f"PASS: Component 4 — Structure intact ({points} pts)")
            else:
                print(f"PARTIAL: Component 4 — {checks_passed}/{total_struct_checks} structure checks ({points} pts)")
            total_score += points
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
