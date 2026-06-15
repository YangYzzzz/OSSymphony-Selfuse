"""
Initial Setup: Email draft with misspellings of 'receive' variants
Task ID: writer_frd_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_039'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Project Status Update - Q1 2025', level=1)

    # --- Email metadata-style header ---
    meta = doc.add_paragraph()
    meta_run = meta.add_run('From: ')
    meta_run.bold = True
    meta.add_run('Sarah Chen <sarah.chen@techvision.com>')
    meta2 = doc.add_paragraph()
    meta2_run = meta2.add_run('To: ')
    meta2_run.bold = True
    meta2.add_run('Project Team <project-team@techvision.com>')
    meta3 = doc.add_paragraph()
    meta3_run = meta3.add_run('Date: ')
    meta3_run.bold = True
    meta3.add_run('March 15, 2025')
    meta4 = doc.add_paragraph()
    meta4_run = meta4.add_run('Subject: ')
    meta4_run.bold = True
    meta4.add_run('Weekly Status Report - Sprint 14')

    doc.add_paragraph('')  # spacer

    # --- Body paragraphs with misspellings ---

    # Paragraph 1: contains 'recieve' (1st occurrence)
    p1 = doc.add_paragraph(
        'Hi Team,'
    )

    p2 = doc.add_paragraph(
        'I wanted to share a quick update on where we stand with the Q1 deliverables. '
        'Please make sure to recieve the attached spreadsheet with the latest metrics '
        'before our meeting on Thursday. The data covers all three product lines and '
        'includes the revised forecasts from the finance department.'
    )

    # Paragraph 2: contains 'recieves' (1st occurrence) and 'recieve' (2nd occurrence)
    p3 = doc.add_paragraph(
        'As a reminder, each team lead recieves a separate breakdown of their '
        "department's budget allocation by end of week. If you did not recieve "
        'last week\'s allocation summary, please reach out to Marcus Johnson in '
        'Finance and he will resend it promptly.'
    )

    # Paragraph 3: contains 'recieved' (1st occurrence)
    p4 = doc.add_paragraph(
        'On the engineering front, we have recieved confirmation from the QA team '
        'that the beta release candidate passed all regression tests. This is a '
        'significant milestone and puts us on track for the April 1st launch date. '
        'The deployment pipeline has been stress-tested and is ready for production.'
    )

    # Paragraph 4: contains 'recieves' (2nd occurrence)
    p5 = doc.add_paragraph(
        'For the client onboarding workflow, the system currently recieves an average '
        'of 45 new registrations per day. Our target is to reach 60 per day by end '
        'of Q2. The marketing team has proposed a new campaign strategy that should '
        'help us close the gap. Details will be shared at the all-hands next Tuesday.'
    )

    # Paragraph 5: contains 'recieve' (3rd occurrence)
    p6 = doc.add_paragraph(
        'Finally, I want to remind everyone that you will recieve your updated access '
        'credentials for the new project management portal by Monday. Please update '
        'your bookmarks and test your login before the old system is deprecated on '
        'March 28th.'
    )

    # Closing
    p7 = doc.add_paragraph(
        'Let me know if you have any questions or concerns. Looking forward to a '
        'productive sprint!'
    )

    doc.add_paragraph('')  # spacer

    p8 = doc.add_paragraph('Best regards,')
    p9 = doc.add_paragraph('Sarah Chen')
    p10 = doc.add_paragraph('Senior Project Manager, TechVision Inc.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
