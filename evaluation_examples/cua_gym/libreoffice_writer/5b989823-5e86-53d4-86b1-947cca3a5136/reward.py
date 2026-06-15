"""
FINAL REWARD SCRIPT - SUCCESS
Task: Every time I hit Enter it just scoots the text down instead of giving me a clean slate. How can I drop in a proper blank page—right after the page I’m working on—without messing up the formatting in the rest of my LibreOffice Writer document?
Generated: 2025-09-10 19:19:34
Status: success
Model: azure-o3
Total Steps: 4
"""

import os
from docx import Document

# ---------------------------------------------------------------
# Reward Verification Script for LibreOffice Writer Blank Page Task
# ---------------------------------------------------------------
# Task (paraphrased):
#   – Insert a proper blank page *after* the current page rather than just    
#     pushing text down with Enter.
#   – The next section/paragraph (“Next Section”) must therefore start on the
#     following page, which is achieved by inserting an actual *page break*   
#     (or equivalent).
#
# Verification Strategy:
#   1. Load the submitted DOCX file (python-docx).
#   2. Detect all explicit page-break elements (w:br w:type="page") that the
#      user inserted.  Each break is associated with the paragraph/run in
#      which it resides; we record the paragraph indices.
#   3. Locate the first paragraph that begins the new section (contains the
#      text "Next Section" – supplied by the golden answer).  In real tasks
#      this could be any marker text defined by the instructions.
#   4. Scoring (progressive):
#        • 0.4 pts  – At least one page break exists in the document.
#        • 0.6 pts  – The identified "Next Section" paragraph occurs *after*
#                     a page break (i.e. starts on the fresh page).
#      If both conditions hold → total = 1.0
#
#   NOTE:  • File existence & successful loading earn *no* points (natural
#            conditions / prerequisites).
#          • The checks are falsifiable – if the break is missing or occurs
#            after the section, points are withheld.
#
# Output: prints detailed diagnostics and final line "REWARD: X.X".
# ---------------------------------------------------------------

def verify_blank_page_task(file_path: str) -> float:
    """Returns a progressive score (0.0-1.0) verifying that a real page break
    precedes the next section, providing a clean blank page."""

    print(f"Verifying document: {file_path}")

    if not os.path.exists(file_path):
        print("✗ File not found – cannot verify task")
        return 0.0  # No points, task failed

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to load DOCX: {e}")
        return 0.0  # Loading failure → task failed

    # -----------------------------------------------------------
    # Step 1: Identify all paragraph indices containing a page break
    # -----------------------------------------------------------
    page_break_indices = []
    for idx, para in enumerate(doc.paragraphs):
        has_break = False
        # Check each run within the paragraph for a <w:br w:type="page"/>
        for run in para.runs:
            for br in run._r.xpath('.//w:br'):
                br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                if br_type == 'page':
                    has_break = True
                    break
            if has_break:
                break
        if has_break:
            page_break_indices.append(idx)

    print(f"Found {len(page_break_indices)} page break paragraph(s): {page_break_indices}")

    # -----------------------------------------------------------
    # Step 2: Locate the start of the next section (marker text)
    # -----------------------------------------------------------
    next_section_index = None
    for idx, para in enumerate(doc.paragraphs):
        if 'Next Section' in para.text:
            next_section_index = idx
            break

    # -----------------------------------------------------------
    # Step 3: Progressive Scoring
    # -----------------------------------------------------------
    score = 0.0

    # 3a. Award for having *any* page break
    if page_break_indices:
        score += 0.4
        print("✓ Page break detected (0.4 pts)")
    else:
        print("✗ No page break detected (0 pts)")

    # 3b. Award additional points if the next section is after a break
    if next_section_index is not None:
        print(f"'Next Section' paragraph found at index {next_section_index}")
        # Determine if at least one break occurs *before* that paragraph
        preceding_break = any(pb_idx < next_section_index for pb_idx in page_break_indices)
        if preceding_break:
            score += 0.6
            print("✓ 'Next Section' starts after a page break (0.6 pts)")
        else:
            print("✗ 'Next Section' does not start after a page break (0 pts)")
    else:
        print("✗ Could not find 'Next Section' paragraph – cannot verify placement (0 pts)")

    final_score = min(score, 1.0)  # Ensure cap at 1.0
    print(f"Total score: {final_score}")
    return final_score

# ---------------------------------------------------------------
# Main Execution – run verification and print REWARD line
# ---------------------------------------------------------------
if __name__ == '__main__':
    FILE_PATH = '/home/user/every_time_i_hit_enter_it_just_scoots_the_text_down_instead_of_giving_me_a_clean_slate_how_can_i_dro.docx'  # <-- Task file path (non-golden)
    reward_value = verify_blank_page_task(FILE_PATH)
    print(f"REWARD: {reward_value}")

