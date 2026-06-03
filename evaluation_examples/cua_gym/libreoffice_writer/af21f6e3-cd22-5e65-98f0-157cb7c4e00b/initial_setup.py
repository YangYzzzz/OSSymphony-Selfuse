"""
Initial Setup: Chemistry notes document with H2O formula (no subscript)
Task ID: writer_txtfmt_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_txtfmt_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP_OUTPUT = f'{WORKDIR}/Desktop/chem_notes.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font for the document via style
    style = doc.styles['Normal']
    style.font.name = 'Liberation Sans'
    style.font.size = Pt(12)

    # Title heading
    title_para = doc.add_heading('Chemistry Notes', level=1)
    for run in title_para.runs:
        run.font.name = 'Liberation Sans'

    # Introductory heading
    section_para = doc.add_heading('Water Molecule', level=2)
    for run in section_para.runs:
        run.font.name = 'Liberation Sans'

    # Main paragraph containing H2O — all text in 12pt Liberation Sans, NO subscript
    # Split into runs to allow the golden patch to easily target the '2'
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run1 = para.add_run('Water (H')
    run1.font.name = 'Liberation Sans'
    run1.font.size = Pt(12)

    run2 = para.add_run('2')
    run2.font.name = 'Liberation Sans'
    run2.font.size = Pt(12)
    run2.font.subscript = False  # explicitly NOT subscript in initial state

    run3 = para.add_run('O) is a polar molecule composed of two hydrogen atoms covalently bonded to a single oxygen atom. Its molecular weight is approximately 18.015 g/mol.')
    run3.font.name = 'Liberation Sans'
    run3.font.size = Pt(12)

    # Additional paragraph for realistic content
    para2 = doc.add_paragraph()
    run4 = para2.add_run('The oxygen atom carries a partial negative charge (\u03b4\u207b) while the two hydrogen atoms carry partial positive charges (\u03b4+), resulting in a dipole moment of 1.85 D.')
    run4.font.name = 'Liberation Sans'
    run4.font.size = Pt(12)

    para3 = doc.add_paragraph()
    run5 = para3.add_run('Water molecules form an extensive hydrogen bond network in liquid state, giving it a high boiling point of 100 \u00b0C at standard pressure.')
    run5.font.name = 'Liberation Sans'
    run5.font.size = Pt(12)

    # Save the main artifact
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also save to Desktop/chem_notes.docx as the task references it there
    import shutil
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)
    shutil.copy(OUTPUT, DESKTOP_OUTPUT)
    print(f'Desktop copy created: {DESKTOP_OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DESKTOP_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
