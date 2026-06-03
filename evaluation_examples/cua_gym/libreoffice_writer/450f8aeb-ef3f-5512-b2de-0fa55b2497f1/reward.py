"""
Reward Script: Employee onboarding checklist with table and checkboxes
Task ID: writer_hr_033
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with 13 rows x 4 columns
  Component 2 (0.25): Header row has correct column names
  Component 3 (0.25): Completed column contains checkbox characters for all 12 task rows
  Component 4 (0.25): All 12 task rows have non-empty Task, Responsible Party, and Due Date
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_033'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table (task requires adding a table)
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document. Task requires an onboarding checklist table.")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table has correct dimensions — 13 rows (header + 12 tasks) x 4 columns (0.25 pts)
    try:
        rows_ok = (num_rows == 13)
        cols_ok = (num_cols == 4)
        if rows_ok and cols_ok:
            print(f"PASS: Component 1 — Table is {num_rows} rows x {num_cols} cols (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Expected 13 rows x 4 cols, found {num_rows} x {num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row has correct column names (0.25 pts)
    try:
        expected_headers = ['completed', 'task', 'responsible party', 'due date']
        actual_headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if actual_headers == expected_headers:
            print(f"PASS: Component 2 — Header row matches expected columns (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Expected headers {expected_headers}, found {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Completed column contains checkbox characters for all 12 task rows (0.25 pts)
    # Checkbox characters include: ☐ (U+2610), ☑ (U+2611), ☒ (U+2612), □ (U+25A1), ■ (U+25A0), ▢, etc.
    try:
        checkbox_chars = {'\u2610', '\u2611', '\u2612', '\u25A1', '\u25A0', '\u25FB', '\u25FC', '\u25FD', '\u25FE'}
        checkbox_count = 0
        for row_idx in range(1, min(num_rows, 13)):
            cell_text = table.rows[row_idx].cells[0].text.strip()
            has_checkbox = any(ch in cell_text for ch in checkbox_chars)
            if has_checkbox:
                checkbox_count += 1
            else:
                print(f"  Row {row_idx} Completed cell: {cell_text!r} — no checkbox char found")

        if checkbox_count == 12:
            print(f"PASS: Component 3 — All 12 task rows have checkbox characters (0.25 pts)")
            total_score += 0.25
        elif checkbox_count >= 8:
            partial = round(0.25 * (checkbox_count / 12), 2)
            print(f"PARTIAL: Component 3 — {checkbox_count}/12 rows have checkboxes ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {checkbox_count}/12 rows have checkbox characters")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: All 12 task rows have non-empty Task, Responsible Party, and Due Date (0.25 pts)
    try:
        complete_rows = 0
        for row_idx in range(1, min(num_rows, 13)):
            task_text = table.rows[row_idx].cells[1].text.strip()
            responsible = table.rows[row_idx].cells[2].text.strip()
            due_date = table.rows[row_idx].cells[3].text.strip()
            if task_text and responsible and due_date:
                complete_rows += 1
            else:
                print(f"  Row {row_idx}: Task={task_text!r}, Responsible={responsible!r}, Due={due_date!r} — incomplete")

        if complete_rows == 12:
            print(f"PASS: Component 4 — All 12 task rows have complete data (0.25 pts)")
            total_score += 0.25
        elif complete_rows >= 8:
            partial = round(0.25 * (complete_rows / 12), 2)
            print(f"PARTIAL: Component 4 — {complete_rows}/12 rows complete ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {complete_rows}/12 rows have complete data")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
