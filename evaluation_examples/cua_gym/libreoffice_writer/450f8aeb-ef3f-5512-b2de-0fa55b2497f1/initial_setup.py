"""
Initial Setup: New Hire Onboarding Checklist - pre-task state
Task ID: writer_hr_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('New Hire Onboarding Checklist', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Human Resources Department')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    run.italic = True

    doc.add_paragraph()  # spacer

    # --- Employee Information Fields ---
    info_heading = doc.add_heading('Employee Information', level=2)

    fields = [
        ('Employee Name', '_' * 40),
        ('Position / Title', '_' * 40),
        ('Department', '_' * 40),
        ('Start Date', '_' * 40),
        ('Direct Manager', '_' * 40),
        ('HR Contact', '_' * 40),
    ]

    for label, placeholder in fields:
        para = doc.add_paragraph()
        run_label = para.add_run(f'{label}: ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_val = para.add_run(placeholder)
        run_val.font.size = Pt(11)
        run_val.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()  # spacer

    # --- Instructions paragraph ---
    instructions = doc.add_paragraph()
    run_inst = instructions.add_run(
        'Please complete all onboarding tasks listed below within the first '
        '30 days of employment. Check off each task as it is completed and '
        'return this form to HR upon completion.'
    )
    run_inst.font.size = Pt(10)
    run_inst.italic = True

    doc.add_paragraph()  # spacer

    # NOTE: The table with checkboxes is NOT included here.
    # That is the task the agent must complete.

    # --- Signature section at bottom ---
    doc.add_paragraph()
    sig_heading = doc.add_heading('Signatures', level=2)

    sig_fields = [
        ('Employee Signature', '_' * 35, 'Date', '_' * 15),
        ('Manager Signature', '_' * 35, 'Date', '_' * 15),
        ('HR Representative', '_' * 35, 'Date', '_' * 15),
    ]

    for label1, val1, label2, val2 in sig_fields:
        para = doc.add_paragraph()
        r1 = para.add_run(f'{label1}: ')
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = para.add_run(f'{val1}    ')
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        r3 = para.add_run(f'{label2}: ')
        r3.bold = True
        r3.font.size = Pt(11)
        r4 = para.add_run(val2)
        r4.font.size = Pt(11)
        r4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
