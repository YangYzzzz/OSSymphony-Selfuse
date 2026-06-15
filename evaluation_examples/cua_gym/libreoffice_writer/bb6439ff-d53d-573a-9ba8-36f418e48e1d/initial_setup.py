"""
Initial Setup: Handbook document with TOC showing chapters 1-7 only (Chapter 8 not yet added via undo)
Task ID: writer_edit_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'handbook'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_entry(doc, chapter_num, chapter_title, page_num):
    """Add a manually-formatted TOC entry paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    # Add tab stop for right-aligned page number
    from docx.shared import Inches
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = para.add_run(f"Chapter {chapter_num}: {chapter_title}")
    run.font.size = Pt(11)
    run2 = para.add_run(f"\t{page_num}")
    run2.font.size = Pt(11)
    return para


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ----- Title Page -----
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_run = title_para.add_run("Employee Handbook")
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run("Human Resources Department")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    edition_para = doc.add_paragraph()
    edition_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    edition_run = edition_para.add_run("Edition 2025 | Revised January 15, 2025")
    edition_run.font.size = Pt(12)

    doc.add_page_break()

    # ----- Table of Contents -----
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.paragraph_format.space_after = Pt(12)

    # TOC entries: Chapters 1-7 only (Chapter 8 is NOT present — undo removed it)
    toc_data = [
        (1, "Introduction to the Company", 3),
        (2, "Code of Conduct", 7),
        (3, "Work Hours and Attendance", 12),
        (4, "Benefits and Compensation", 18),
        (5, "Leave Policies", 25),
        (6, "Performance and Development", 31),
        (7, "Health and Safety", 38),
    ]

    for num, title, page in toc_data:
        add_toc_entry(doc, num, title, page)

    doc.add_page_break()

    # ----- Chapter 1 -----
    doc.add_heading('Chapter 1: Introduction to the Company', level=1)
    doc.add_paragraph(
        "Welcome to Meridian Solutions Group. We are a globally recognized technology "
        "and professional services firm committed to innovation, integrity, and inclusion. "
        "Founded in 2003, Meridian Solutions Group has grown from a small consulting boutique "
        "to a multinational organization with offices in 28 countries and a workforce of over "
        "12,000 dedicated professionals."
    )
    doc.add_paragraph(
        "Our mission is to deliver transformational solutions that empower our clients to "
        "achieve their strategic objectives. We pride ourselves on a culture of continuous "
        "improvement and collaborative problem-solving. This handbook has been designed to "
        "help you understand our values, policies, and the support systems available to you."
    )
    doc.add_paragraph(
        "As a new or continuing member of our team, you play a critical role in shaping the "
        "future of Meridian Solutions Group. We encourage you to read this handbook carefully "
        "and reach out to your HR Business Partner if you have any questions."
    )

    doc.add_heading('1.1 Company History', level=2)
    doc.add_paragraph(
        "Meridian Solutions Group was established by founders Dr. Eleanor Whitfield and "
        "Mr. James Ramirez in San Francisco, California. The company initially focused on "
        "financial technology advisory services before expanding into enterprise software, "
        "cloud infrastructure, and digital transformation consulting."
    )

    doc.add_page_break()

    # ----- Chapter 2 -----
    doc.add_heading('Chapter 2: Code of Conduct', level=1)
    doc.add_paragraph(
        "All employees of Meridian Solutions Group are expected to uphold the highest "
        "standards of professional behavior. Our Code of Conduct provides the framework "
        "for ethical decision-making and outlines the expectations for how we treat one "
        "another, our clients, and our broader community."
    )
    doc.add_paragraph(
        "Key principles include honesty in all dealings, respect for colleagues regardless "
        "of background or seniority, protection of confidential information, and avoidance "
        "of conflicts of interest. Violations of the Code of Conduct may result in "
        "disciplinary action up to and including termination."
    )

    doc.add_heading('2.1 Workplace Harassment Policy', level=2)
    doc.add_paragraph(
        "Meridian Solutions Group maintains a zero-tolerance policy for harassment of any "
        "kind. This includes verbal, physical, and written conduct that creates a hostile "
        "work environment. All reports of harassment are treated with strict confidentiality "
        "and investigated thoroughly by the People & Culture team."
    )

    doc.add_page_break()

    # ----- Chapter 3 -----
    doc.add_heading('Chapter 3: Work Hours and Attendance', level=1)
    doc.add_paragraph(
        "Standard working hours at Meridian Solutions Group are Monday through Friday, "
        "9:00 AM to 5:30 PM local time, totaling 40 hours per week. Flexible work "
        "arrangements are available upon manager approval and subject to business needs."
    )
    doc.add_paragraph(
        "Employees are expected to maintain reliable attendance and to notify their "
        "direct manager and HR no later than 8:00 AM on days when they are unable to "
        "attend work. Repeated unexcused absences may be subject to formal review."
    )

    doc.add_page_break()

    # ----- Chapter 4 -----
    doc.add_heading('Chapter 4: Benefits and Compensation', level=1)
    doc.add_paragraph(
        "Meridian Solutions Group offers a comprehensive benefits package designed to "
        "support the health, financial security, and well-being of all full-time employees. "
        "Benefits are reviewed annually and may be updated to reflect market conditions "
        "and employee feedback."
    )

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    headers = ['Benefit', 'Coverage Details']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(11)

    benefit_rows = [
        ('Medical Insurance', 'Full coverage for employee; 80% for dependents'),
        ('Dental & Vision', 'Up to $2,500 per annum combined'),
        ('401(k) Retirement', '5% company match on employee contributions'),
        ('Life Insurance', '2x annual salary, at no cost to employee'),
        ('Wellness Stipend', '$600 per year for gym, fitness, or wellness apps'),
    ]
    for row_idx, (benefit, details) in enumerate(benefit_rows, 1):
        table.cell(row_idx, 0).text = benefit
        table.cell(row_idx, 1).text = details

    doc.add_page_break()

    # ----- Chapter 5 -----
    doc.add_heading('Chapter 5: Leave Policies', level=1)
    doc.add_paragraph(
        "Meridian Solutions Group provides a generous leave policy that includes annual "
        "leave, sick leave, parental leave, and compassionate leave. Employees are "
        "encouraged to take their entitled leave to support work-life balance."
    )
    doc.add_paragraph(
        "Annual leave accrues at a rate of 1.67 days per month for full-time employees, "
        "totaling 20 days per year. Unused annual leave of up to 10 days may be carried "
        "forward to the next calendar year. Sick leave is provided at 10 days per year "
        "and does not accrue."
    )

    doc.add_page_break()

    # ----- Chapter 6 -----
    doc.add_heading('Chapter 6: Performance and Development', level=1)
    doc.add_paragraph(
        "Our performance management process is designed to support continuous growth "
        "and recognize exceptional contributions. Annual performance reviews take place "
        "in December, with mid-year check-ins in June."
    )
    doc.add_paragraph(
        "Each employee works with their manager to set SMART objectives at the start of "
        "each year. These objectives form the basis of both mid-year and end-of-year "
        "evaluations. The performance rating scale ranges from 1 (Needs Improvement) "
        "to 5 (Exceptional Performer)."
    )

    doc.add_heading('6.1 Learning & Development Budget', level=2)
    doc.add_paragraph(
        "All employees are entitled to an annual Learning & Development budget of $1,500 "
        "for approved training, certifications, or conferences. Applications must be "
        "submitted through the HR portal with manager approval at least 30 days in advance."
    )

    doc.add_page_break()

    # ----- Chapter 7 -----
    doc.add_heading('Chapter 7: Health and Safety', level=1)
    doc.add_paragraph(
        "The health and safety of all employees is a top priority at Meridian Solutions Group. "
        "We comply with all applicable occupational health and safety regulations and "
        "maintain a proactive safety culture across all offices and remote work environments."
    )
    doc.add_paragraph(
        "Emergency procedures are posted at all office entrances and reviewed during "
        "new employee onboarding. Fire drills are conducted twice per year. Employees "
        "are required to report any workplace hazards or near-miss incidents to the "
        "Facilities team immediately."
    )

    doc.add_heading('7.1 Ergonomics and Remote Work', level=2)
    doc.add_paragraph(
        "Employees working from home are entitled to a one-time home office setup allowance "
        "of $800. This may be used for ergonomic chairs, standing desks, monitors, keyboards, "
        "or other approved equipment. Receipts must be submitted via the expense portal within "
        "60 days of purchase."
    )

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
