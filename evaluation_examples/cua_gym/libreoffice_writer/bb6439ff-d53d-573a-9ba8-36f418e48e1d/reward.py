"""
Reward Script: Redo action to restore 'Chapter 8: Troubleshooting' in Table of Contents
Task ID: writer_edit_046
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): 'Chapter 8: Troubleshooting' entry exists in the Table of Contents section
  Component 2 (0.3 pts): 'Chapter 8: Troubleshooting' heading exists in the document body
  Component 3 (0.2 pts): Chapter 8 TOC entry includes a page number reference (tab + digit)
Total: 1.0
"""

import os
import re

from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_edit_046'
FILE_PATH = f'{WORKDIR}/handbook.docx'


def verify_task(file_path):
    """
    Verify that the Redo operation was performed:
    - The TOC should contain 'Chapter 8: Troubleshooting' (the re-done change)
    - The document body should contain the Chapter 8 heading
    - The TOC entry should include a page number after the tab
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the Table of Contents section and its entries
    # Locate the TOC heading index, then collect paragraphs until the next Heading 1
    all_paras = doc.paragraphs
    toc_start_idx = next(
        (i for i, p in enumerate(all_paras)
         if p.style.name == 'Heading 1' and 'table of contents' in p.text.strip().lower()),
        None
    )
    toc_paragraphs = []
    if toc_start_idx is not None:
        for para in all_paras[toc_start_idx + 1:]:
            if para.style.name == 'Heading 1':
                break
            toc_paragraphs.append(para)

    toc_text_all = '\n'.join(p.text for p in toc_paragraphs)

    # Component 1: 'Chapter 8: Troubleshooting' exists in the Table of Contents (0.5 points)
    # This FAILS on initial_env (no Chapter 8 TOC entry) and PASSES on golden_env
    try:
        ch8_in_toc = any(
            'chapter 8' in p.text.lower() and 'troubleshooting' in p.text.lower()
            for p in toc_paragraphs
        )
        if ch8_in_toc:
            ch8_toc_text = next(
                p.text for p in toc_paragraphs
                if 'chapter 8' in p.text.lower() and 'troubleshooting' in p.text.lower()
            )
            print(f"PASS: Component 1 — 'Chapter 8: Troubleshooting' found in TOC: {repr(ch8_toc_text)} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — 'Chapter 8: Troubleshooting' NOT found in TOC. TOC entries: {repr(toc_text_all)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 'Chapter 8: Troubleshooting' heading exists in the document body (0.3 points)
    # This FAILS on initial_env (no Chapter 8 section) and PASSES on golden_env
    try:
        ch8_heading_found = any(
            para.style.name in ('Heading 1', 'Heading 2')
            and 'chapter 8' in para.text.lower()
            and 'troubleshooting' in para.text.lower()
            for para in doc.paragraphs
        )
        if ch8_heading_found:
            print(f"PASS: Component 2 — 'Chapter 8: Troubleshooting' heading found in document body (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — 'Chapter 8: Troubleshooting' heading NOT found in document body")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Chapter 8 TOC entry includes a page number (tab + digits) (0.2 points)
    # This FAILS on initial_env (no entry exists) and PASSES on golden_env
    try:
        ch8_toc_entry = next(
            (p.text for p in toc_paragraphs
             if 'chapter 8' in p.text.lower() and 'troubleshooting' in p.text.lower()),
            None
        )
        if ch8_toc_entry is not None and re.search(r'\t\d+', ch8_toc_entry):
            page_num = re.search(r'\t(\d+)', ch8_toc_entry).group(1)
            print(f"PASS: Component 3 — Chapter 8 TOC entry has page number reference: {repr(ch8_toc_entry)} (0.2 pts)")
            total_score += 0.2
        else:
            if ch8_toc_entry is None:
                print(f"FAIL: Component 3 — No Chapter 8 TOC entry found (prerequisite for page number check)")
            else:
                print(f"FAIL: Component 3 — Chapter 8 TOC entry found but missing page number: {repr(ch8_toc_entry)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
