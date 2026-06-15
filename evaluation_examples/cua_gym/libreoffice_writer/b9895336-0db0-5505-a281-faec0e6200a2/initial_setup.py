"""
Initial Setup: Create a formatted NDA document for template conversion task
Task ID: writer_legal_080
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_080'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Header (Firm Letterhead) ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_firm = hp.add_run("WHITFIELD, TORRES & ASSOCIATES LLP")
    run_firm.bold = True
    run_firm.font.size = Pt(14)
    run_firm.font.name = "Times New Roman"
    run_firm.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    hp2 = header.add_paragraph()
    hp2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_addr = hp2.add_run("1200 Pacific Avenue, Suite 4500  |  San Francisco, CA 94105")
    run_addr.font.size = Pt(9)
    run_addr.font.name = "Times New Roman"
    run_addr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    hp3 = header.add_paragraph()
    hp3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_contact = hp3.add_run("Tel: (415) 555-0198  |  Fax: (415) 555-0199  |  www.whitfieldtorres.com")
    run_contact.font.size = Pt(9)
    run_contact.font.name = "Times New Roman"
    run_contact.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Footer ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_footer = fp.add_run("CONFIDENTIAL - Attorney-Client Privileged")
    run_footer.font.size = Pt(8)
    run_footer.font.name = "Times New Roman"
    run_footer.font.italic = True
    run_footer.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(6)
    run_title = title_para.add_run("NON-DISCLOSURE AGREEMENT")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run_sub = subtitle.add_run("(Mutual Confidentiality Agreement)")
    run_sub.font.size = Pt(11)
    run_sub.font.name = "Times New Roman"
    run_sub.font.italic = True

    # Helper function for adding styled paragraphs
    def add_heading_text(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"
        return p

    def add_body_text(text, space_after=Pt(6)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = space_after
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        return p

    # --- Preamble ---
    add_body_text(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of '
        '__________________ ("Effective Date") by and between:'
    )

    add_body_text(
        'Party A: ______________________________ ("Disclosing Party"), '
        'a corporation organized under the laws of the State of California, '
        'with its principal place of business at ______________________________; and'
    )

    add_body_text(
        'Party B: ______________________________ ("Receiving Party"), '
        'a corporation organized under the laws of the State of ______________, '
        'with its principal place of business at ______________________________.'
    )

    add_body_text(
        'The Disclosing Party and the Receiving Party are hereinafter collectively '
        'referred to as the "Parties" and individually as a "Party."'
    )

    # --- RECITALS ---
    add_heading_text("RECITALS")
    add_body_text(
        "WHEREAS, the Parties wish to explore a potential business relationship "
        '("Purpose") and in connection therewith, each Party may disclose certain '
        "confidential and proprietary information to the other Party; and"
    )
    add_body_text(
        "WHEREAS, the Parties desire to establish the terms and conditions under "
        "which such confidential information will be disclosed and protected."
    )
    add_body_text(
        "NOW, THEREFORE, in consideration of the mutual covenants and agreements "
        "set forth herein, and for other good and valuable consideration, the "
        "receipt and sufficiency of which are hereby acknowledged, the Parties "
        "agree as follows:"
    )

    # --- Section 1: Definitions ---
    add_heading_text("1. DEFINITIONS")
    add_body_text(
        '1.1 "Confidential Information" means any and all non-public, proprietary, '
        "or confidential information disclosed by either Party to the other Party, "
        "whether orally, in writing, electronically, or by any other means, including "
        "but not limited to: (a) trade secrets, inventions, patents, copyrights, "
        "trademarks, and other intellectual property; (b) business plans, strategies, "
        "financial data, projections, and operational information; (c) customer lists, "
        "vendor relationships, and market analyses; (d) technical data, designs, "
        "specifications, algorithms, and source code; and (e) any information that a "
        "reasonable person would understand to be confidential given the nature of "
        "the information and the circumstances of disclosure."
    )
    add_body_text(
        '1.2 "Confidential Information" does not include information that: '
        "(a) is or becomes publicly available through no fault of the Receiving Party; "
        "(b) was already known to the Receiving Party prior to disclosure, as evidenced "
        "by written records; (c) is independently developed by the Receiving Party "
        "without use of or reference to the Confidential Information; or (d) is "
        "lawfully obtained from a third party without restriction on disclosure."
    )

    # --- Section 2: Obligations ---
    add_heading_text("2. OBLIGATIONS OF THE RECEIVING PARTY")
    add_body_text(
        "2.1 The Receiving Party shall: (a) hold and maintain the Confidential "
        "Information in strict confidence using at least the same degree of care as "
        "it uses to protect its own confidential information, but in no event less "
        "than a reasonable degree of care; (b) not disclose any Confidential "
        "Information to any third parties without the prior written consent of the "
        "Disclosing Party; (c) not use the Confidential Information for any purpose "
        "other than the Purpose; and (d) limit access to the Confidential Information "
        "to those of its employees, agents, and advisors who have a need to know and "
        "who are bound by confidentiality obligations at least as restrictive as those "
        "contained herein."
    )
    add_body_text(
        "2.2 The Receiving Party shall promptly notify the Disclosing Party in "
        "writing of any unauthorized use or disclosure of Confidential Information "
        "and shall cooperate with the Disclosing Party to prevent further "
        "unauthorized use or disclosure."
    )

    # --- Section 3: Term ---
    add_heading_text("3. TERM AND TERMINATION")
    add_body_text(
        "3.1 This Agreement shall be effective as of the Effective Date and shall "
        "continue in full force and effect for a period of three (3) years from the "
        "Effective Date, unless earlier terminated by either Party upon thirty (30) "
        "days' prior written notice to the other Party."
    )
    add_body_text(
        "3.2 The obligations of confidentiality set forth herein shall survive the "
        "termination or expiration of this Agreement for a period of five (5) years "
        "following such termination or expiration."
    )

    # --- Section 4: Return of Materials ---
    add_heading_text("4. RETURN OF MATERIALS")
    add_body_text(
        "4.1 Upon termination of this Agreement or upon written request by the "
        "Disclosing Party, the Receiving Party shall promptly: (a) return all "
        "tangible materials containing Confidential Information; (b) destroy all "
        "copies, notes, summaries, and other reproductions of Confidential "
        "Information in any form; and (c) certify in writing that all such "
        "materials have been returned or destroyed."
    )

    # --- Section 5: Remedies ---
    add_heading_text("5. REMEDIES")
    add_body_text(
        "5.1 The Parties acknowledge that any breach of this Agreement may cause "
        "irreparable harm to the Disclosing Party, for which monetary damages may "
        "be inadequate. Accordingly, the Disclosing Party shall be entitled to seek "
        "equitable relief, including injunction and specific performance, in addition "
        "to all other remedies available at law or in equity."
    )

    # --- Section 6: General Provisions ---
    add_heading_text("6. GENERAL PROVISIONS")
    add_body_text(
        "6.1 Governing Law. This Agreement shall be governed by and construed in "
        "accordance with the laws of the State of California, without regard to its "
        "conflict of laws principles."
    )
    add_body_text(
        "6.2 Entire Agreement. This Agreement constitutes the entire agreement "
        "between the Parties with respect to the subject matter hereof and supersedes "
        "all prior and contemporaneous agreements, understandings, negotiations, and "
        "discussions, whether oral or written."
    )
    add_body_text(
        "6.3 Amendment. This Agreement may not be amended or modified except by a "
        "written instrument signed by both Parties."
    )
    add_body_text(
        "6.4 Severability. If any provision of this Agreement is held to be invalid "
        "or unenforceable, the remaining provisions shall continue in full force and "
        "effect."
    )
    add_body_text(
        "6.5 Waiver. The failure of either Party to enforce any provision of this "
        "Agreement shall not constitute a waiver of such provision or the right to "
        "enforce it at a later time."
    )
    add_body_text(
        "6.6 Assignment. Neither Party may assign or transfer this Agreement without "
        "the prior written consent of the other Party."
    )

    # --- Signature Blocks ---
    add_heading_text("")  # spacer
    sig_line = "________________________________________"

    add_body_text("IN WITNESS WHEREOF, the Parties have executed this Non-Disclosure "
                  "Agreement as of the date first written above.", space_after=Pt(24))

    # Disclosing Party signature block
    add_heading_text("DISCLOSING PARTY:")
    add_body_text("")
    add_body_text(sig_line)
    add_body_text("Name: ______________________________")
    add_body_text("Title: ______________________________")
    add_body_text("Date: ______________________________", space_after=Pt(18))

    # Receiving Party signature block
    add_heading_text("RECEIVING PARTY:")
    add_body_text("")
    add_body_text(sig_line)
    add_body_text("Name: ______________________________")
    add_body_text("Title: ______________________________")
    add_body_text("Date: ______________________________")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
