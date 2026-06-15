"""
Initial Setup: Business analysis report with bibliography section (no Johnson 2022 entry)
Task ID: osworld_writer_bibliography_crossref_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Corporate Governance and Business Ethics: An Analysis', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introduction / Body Paragraph 1 ---
    p1 = doc.add_paragraph(
        'Corporate governance has emerged as a critical framework for ensuring accountability, '
        'transparency, and ethical conduct within modern organizations. As stakeholders demand '
        'greater oversight of corporate activities, companies are increasingly pressured to adopt '
        'robust governance structures that align executive decisions with long-term shareholder value. '
        'Smith and colleagues (2019) argue that effective governance mechanisms reduce agency costs '
        'and improve overall firm performance across diverse industry sectors.'
    )

    # --- Body Paragraph 2 (citation will be inserted here at the end) ---
    p2 = doc.add_paragraph(
        'The relationship between corporate governance and financial performance has been extensively '
        'studied over the past two decades. Board composition, executive compensation, and audit '
        'committee independence are among the most frequently examined governance attributes. '
        'Research consistently demonstrates that firms with stronger governance practices tend to '
        'exhibit lower volatility, higher return on equity, and greater investor confidence. '
        'Furthermore, governance quality has been linked to improved strategic decision-making '
        'and reduced instances of financial misconduct within publicly listed companies.'
    )

    # --- Body Paragraph 3 ---
    p3 = doc.add_paragraph(
        'Emerging markets present unique challenges for corporate governance implementation, '
        'as institutional frameworks and regulatory environments differ significantly from '
        'those in developed economies. Cultural norms, ownership concentration, and weak '
        'enforcement mechanisms can undermine even well-designed governance codes. '
        'Patel and Nguyen (2021) highlight the importance of adapting governance frameworks '
        'to local institutional contexts rather than adopting a one-size-fits-all approach. '
        'International organizations and development agencies continue to promote best '
        'practices in governance as a tool for sustainable economic development.'
    )

    # --- Bibliography Section ---
    doc.add_heading('Bibliography', level=1)

    # Entry 1
    bib1 = doc.add_paragraph(
        'Patel, R., & Nguyen, T. (2021). Governance in Emerging Markets: Challenges and Adaptations. '
        'International Business Review, 30(4), 101-118.'
    )
    bib1.paragraph_format.left_indent = Pt(0)

    # Entry 2
    bib2 = doc.add_paragraph(
        'Smith, A., Brown, C., & Davis, E. (2019). Agency Theory and Corporate Governance: '
        'A Meta-Analysis. Journal of Corporate Finance, 55(2), 45-67.'
    )
    bib2.paragraph_format.left_indent = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
