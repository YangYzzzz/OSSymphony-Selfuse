"""
Reward Script: Add bibliography reference and in-text citation for Johnson (2022)
Task ID: osworld_writer_bibliography_crossref_003
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Citation '(Johnson, 2022)' appended at end of second paragraph
  Component 2 (0.4): Bibliography entry 'Johnson, M. (2022)...' added after existing bibliography entries
  Component 3 (0.2): Bibliography entry contains the key required fields (author, year, title, publisher)
  Total: 1.0
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_bibliography_crossref_003'


def find_bib_heading_idx(doc):
    """Return the paragraph index of the Bibliography heading, or None."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and 'bibliography' in p.text.lower():
            return i
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.

    Task: Add reference 'Johnson, M. (2022). Corporate Governance Frameworks. Business Press.'
    to the bibliography and insert citation '(Johnson, 2022)' at the end of the second paragraph.

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: Citation '(Johnson, 2022)' appended at the end of the second body paragraph (0.4 points) ---
    try:
        bib_heading_idx = find_bib_heading_idx(doc)

        # Collect body paragraphs before the bibliography heading
        if bib_heading_idx is not None:
            body_only_paras = [p for p in doc.paragraphs[:bib_heading_idx]
                               if p.text.strip() and p.style.name not in ('Title',)]
        else:
            # Fallback: non-empty paragraphs excluding headings/title
            body_only_paras = [p for p in doc.paragraphs
                               if p.text.strip() and p.style.name not in ('Title', 'Heading 1',
                                                                           'Heading 2', 'Heading 3')]

        if len(body_only_paras) >= 2:
            second_para = body_only_paras[1]
            second_para_text = second_para.text.strip()
            # Check that the citation appears at the end of the second paragraph
            if re.search(r'\(Johnson,?\s*2022\)\s*$', second_para_text):
                print(f"PASS: Component 1 — Citation '(Johnson, 2022)' found at end of second paragraph (0.4 pts)")
                print(f"      Second para ends with: ...{second_para_text[-50:]!r}")
                total_score += 0.4
            elif '(Johnson, 2022)' in second_para_text or '(Johnson,2022)' in second_para_text:
                # Citation present but not at end — partial credit
                print(f"PARTIAL: Component 1 — Citation '(Johnson, 2022)' found in second paragraph but NOT at end (0.2 pts)")
                print(f"         Second para: ...{second_para_text[-100:]!r}")
                total_score += 0.2
            else:
                print(f"FAIL: Component 1 — Citation '(Johnson, 2022)' not found at end of second paragraph")
                print(f"      Second para ends with: ...{second_para_text[-80:]!r}")
        else:
            print(f"FAIL: Component 1 — Could not identify second body paragraph (found {len(body_only_paras)} body paragraphs)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: Bibliography entry for Johnson (2022) present after existing entries (0.4 points) ---
    try:
        bib_heading_idx = find_bib_heading_idx(doc)

        if bib_heading_idx is not None:
            bib_entries = [p for p in doc.paragraphs[bib_heading_idx+1:] if p.text.strip()]
        else:
            bib_entries = []
            print("WARN: Component 2 — Could not find Bibliography heading")

        # Find Johnson entry in bibliography
        johnson_entry_text = ""
        for entry in bib_entries:
            if re.search(r'Johnson.*2022', entry.text.strip()):
                johnson_entry_text = entry.text.strip()
                break

        if johnson_entry_text:
            print(f"PASS: Component 2 — Johnson (2022) bibliography entry found (0.4 pts)")
            print(f"      Entry: {johnson_entry_text!r}")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — No Johnson (2022) entry in bibliography")
            if bib_entries:
                print(f"      Existing entries: {[e.text[:60] for e in bib_entries]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: Bibliography entry contains required bibliographic fields (0.2 points) ---
    # Required: Author (Johnson, M.), Year (2022), Title (Corporate Governance Frameworks), Publisher (Business Press)
    try:
        bib_heading_idx = find_bib_heading_idx(doc)
        johnson_full_text = ""

        if bib_heading_idx is not None:
            for p in doc.paragraphs[bib_heading_idx+1:]:
                if 'Johnson' in p.text and '2022' in p.text and p.text.strip():
                    johnson_full_text = p.text.strip()
                    break

        if johnson_full_text:
            # Check required elements
            checks = {
                'author': bool(re.search(r'Johnson,?\s*M\.?', johnson_full_text)),
                'year': '2022' in johnson_full_text,
                'title': 'Corporate Governance Frameworks' in johnson_full_text,
                'publisher': 'Business Press' in johnson_full_text,
            }
            passed = sum(checks.values())
            if passed == 4:
                print(f"PASS: Component 3 — All bibliography fields present: author, year, title, publisher (0.2 pts)")
                total_score += 0.2
            elif passed >= 2:
                print(f"PARTIAL: Component 3 — {passed}/4 bibliography fields present: {checks} (0.1 pts)")
                total_score += 0.1
            else:
                print(f"FAIL: Component 3 — Missing key bibliography fields: {checks}")
                print(f"      Entry text: {johnson_full_text!r}")
        else:
            print(f"FAIL: Component 3 — No Johnson (2022) entry found in bibliography for field verification")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
