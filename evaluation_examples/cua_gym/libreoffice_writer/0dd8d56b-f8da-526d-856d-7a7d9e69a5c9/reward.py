"""
Reward Script: Patient intake form for Riverside Medical Clinic
Task ID: writer_wf_054
Domain: libreoffice_writer
Scoring:
  C1: Title bold+centered (0.15)
  C2: Personal Information table 7 fields (0.20)
  C3: Insurance Information table 3 fields (0.15)
  C4: Medical History 8 checkbox conditions (0.20)
  C5: Medications table header+4 rows (0.15)
  C6: Consent/Signature section (0.15)
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_054'


def persist_app_state(domain):
    """Try to save any unsaved changes in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify patient intake form with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_text = ' '.join(p.text for p in doc.paragraphs).lower()
    tables = doc.tables

    # Component 1: Title "Riverside Medical Clinic" bold and centered (0.15 points)
    try:
        found_title = False
        for para in doc.paragraphs:
            if 'riverside' in para.text.lower() and 'medical' in para.text.lower() and 'clinic' in para.text.lower():
                is_bold = any(r.bold for r in para.runs if r.text.strip())
                is_centered = para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                if is_bold and is_centered:
                    found_title = True
                    print(f"PASS: Component 1 — Title '{para.text}' is bold and centered (0.15 pts)")
                    total_score += 0.15
                elif is_bold:
                    print(f"PARTIAL: Component 1 — Title is bold but not centered")
                    total_score += 0.07
                elif is_centered:
                    print(f"PARTIAL: Component 1 — Title is centered but not bold")
                    total_score += 0.07
                else:
                    print(f"FAIL: Component 1 — Title found but not bold or centered")
                break
        if not found_title and total_score == 0.0:
            print("FAIL: Component 1 — No 'Riverside Medical Clinic' title found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Personal Information table with 7 fields (0.20 points)
    # Expected fields: Name, DOB/Date of Birth, Gender, Address, Phone, Email, Emergency Contact
    try:
        personal_fields = ['name', 'date of birth', 'dob', 'gender', 'address', 'phone', 'email', 'emergency contact']
        # Find a table that has personal info fields
        best_match_count = 0
        for table in tables:
            matched = set()
            for row in table.rows:
                cell_text = row.cells[0].text.strip().lower()
                for field in personal_fields:
                    if field in cell_text:
                        matched.add(field)
            # DOB and Date of Birth are the same field
            if 'dob' in matched or 'date of birth' in matched:
                matched.add('dob')
                matched.add('date of birth')
            # Count unique logical fields (7 total)
            logical = set()
            for m in matched:
                if m in ('dob', 'date of birth'):
                    logical.add('dob')
                else:
                    logical.add(m)
            best_match_count = max(best_match_count, len(logical))

        if best_match_count >= 7:
            print(f"PASS: Component 2 — Personal Information table has {best_match_count}/7 fields (0.20 pts)")
            total_score += 0.20
        elif best_match_count >= 5:
            pts = round(0.20 * (best_match_count / 7), 2)
            print(f"PARTIAL: Component 2 — Personal Information table has {best_match_count}/7 fields ({pts} pts)")
            total_score += pts
        elif best_match_count >= 3:
            pts = round(0.20 * (best_match_count / 7), 2)
            print(f"PARTIAL: Component 2 — Personal Information table has {best_match_count}/7 fields ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 2 — Personal Information table not found or has only {best_match_count}/7 fields")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Insurance Information table with 3 fields (0.15 points)
    # Expected: Provider, Policy #, Group #
    try:
        insurance_fields = ['provider', 'policy', 'group']
        best_ins_count = 0
        for table in tables:
            matched = set()
            for row in table.rows:
                cell_text = row.cells[0].text.strip().lower()
                for field in insurance_fields:
                    if field in cell_text:
                        matched.add(field)
            best_ins_count = max(best_ins_count, len(matched))

        if best_ins_count >= 3:
            print(f"PASS: Component 3 — Insurance Information table has {best_ins_count}/3 fields (0.15 pts)")
            total_score += 0.15
        elif best_ins_count >= 2:
            pts = round(0.15 * (best_ins_count / 3), 2)
            print(f"PARTIAL: Component 3 — Insurance table has {best_ins_count}/3 fields ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 3 — Insurance Information table not found or has only {best_ins_count}/3 fields")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Medical History with 8 checkbox conditions (0.20 points)
    # Expected: 8 conditions with checkbox character (☐ or similar)
    try:
        checkbox_chars = ['\u2610', '\u2611', '\u2612', '\u25a1', '\u25a0', '[ ]', '[x]', '[X]', '[]']
        conditions = ['diabetes', 'heart disease', 'high blood pressure', 'hypertension',
                       'asthma', 'cancer', 'arthritis', 'depression', 'anxiety', 'allergies']
        found_conditions = set()
        has_checkboxes = False

        for para in doc.paragraphs:
            text_lower = para.text.lower().strip()
            # Check if this looks like a checkbox item
            has_cb = any(cb in para.text for cb in checkbox_chars)
            for cond in conditions:
                if cond in text_lower:
                    found_conditions.add(cond)
                    if has_cb:
                        has_checkboxes = True

        # depression and anxiety may be combined
        logical_count = len(found_conditions)
        # Combine depression/anxiety as one
        if 'depression' in found_conditions and 'anxiety' in found_conditions:
            logical_count -= 1
        # hypertension and high blood pressure are the same
        if 'hypertension' in found_conditions and 'high blood pressure' in found_conditions:
            logical_count -= 1

        if logical_count >= 8 and has_checkboxes:
            print(f"PASS: Component 4 — Medical History has {logical_count} conditions with checkboxes (0.20 pts)")
            total_score += 0.20
        elif logical_count >= 6:
            pts = round(0.20 * min(logical_count, 8) / 8, 2)
            if not has_checkboxes:
                pts = round(pts * 0.5, 2)  # penalize missing checkboxes
            print(f"PARTIAL: Component 4 — Medical History has {logical_count} conditions, checkboxes={'yes' if has_checkboxes else 'no'} ({pts} pts)")
            total_score += pts
        elif logical_count >= 3:
            pts = round(0.20 * logical_count / 8, 2)
            print(f"PARTIAL: Component 4 — Medical History has {logical_count}/8 conditions ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 4 — Medical History has only {logical_count} conditions")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Current Medications table with header + 4 rows (0.15 points)
    # Expected: table with columns Medication, Dosage, Frequency and at least 4 data rows
    try:
        found_med_table = False
        for table in tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            has_med = any('medication' in c or 'med' == c for c in header_cells)
            has_dosage = any('dosage' in c or 'dose' in c for c in header_cells)
            has_freq = any('frequency' in c or 'freq' in c for c in header_cells)

            if has_med and has_dosage and has_freq:
                found_med_table = True
                data_rows = len(table.rows) - 1  # exclude header
                if data_rows >= 4:
                    print(f"PASS: Component 5 — Medications table has header + {data_rows} rows (0.15 pts)")
                    total_score += 0.15
                elif data_rows >= 2:
                    pts = round(0.15 * data_rows / 4, 2)
                    print(f"PARTIAL: Component 5 — Medications table has header + {data_rows} rows ({pts} pts)")
                    total_score += pts
                else:
                    print(f"FAIL: Component 5 — Medications table found but only {data_rows} data rows")
                break

        if not found_med_table:
            print("FAIL: Component 5 — No Medications table found with Medication/Dosage/Frequency columns")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Consent and Signature section (0.15 points)
    # Expected: consent text + signature line + date line
    try:
        has_consent_text = False
        has_signature = False
        has_date_line = False

        for para in doc.paragraphs:
            text_lower = para.text.lower()
            if 'consent' in text_lower and ('personal' in text_lower or 'health' in text_lower or 'information' in text_lower or 'agree' in text_lower or 'authorize' in text_lower or 'hereby' in text_lower):
                has_consent_text = True
            if 'signature' in text_lower and ('_' in para.text or 'sign' in text_lower):
                has_signature = True
            if 'date' in text_lower and ('_' in para.text or ':' in para.text):
                has_date_line = True

        parts = sum([has_consent_text, has_signature, has_date_line])
        if parts >= 3:
            print(f"PASS: Component 6 — Consent section has consent text, signature line, and date line (0.15 pts)")
            total_score += 0.15
        elif parts >= 2:
            pts = round(0.15 * parts / 3, 2)
            print(f"PARTIAL: Component 6 — Consent section has {parts}/3 elements ({pts} pts)")
            total_score += pts
        elif parts >= 1:
            pts = round(0.15 * parts / 3, 2)
            print(f"PARTIAL: Component 6 — Consent section has {parts}/3 elements ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 6 — No consent/signature section found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
