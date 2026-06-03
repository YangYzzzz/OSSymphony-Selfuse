"""
Reward Script: Create thesis chapter document with TOC, headings, footnotes, and bibliography
Task ID: writer_wf_002
Domain: libreoffice_writer
Scoring:
  Component 1: Chapter title 'Chapter 3: Methodology' as Heading 1 (0.15)
  Component 2: Four Heading 2 sections present (0.25)
  Component 3: At least 2 footnotes (0.20)
  Component 4: Bibliography section with 3+ references (0.20)
  Component 5: Table of Contents present at beginning (0.20)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_002'


def persist_app_state(domain: str):
    """Best-effort save any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all headings for reuse
    headings_h1 = []
    headings_h2 = []
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Heading 1':
            headings_h1.append(p.text.strip())
        elif p.style and p.style.name == 'Heading 2':
            headings_h2.append(p.text.strip())

    # Component 1: Chapter title 'Chapter 3: Methodology' as Heading 1 (0.15 points)
    try:
        chapter_title_found = any(
            'chapter 3' in h.lower() and 'methodology' in h.lower()
            for h in headings_h1
        )
        if chapter_title_found:
            print(f"PASS: Component 1 — 'Chapter 3: Methodology' found as Heading 1 (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — No Heading 1 matching 'Chapter 3: Methodology'. "
                  f"Found H1s: {headings_h1}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Four required Heading 2 sections (0.25 points — 0.0625 each)
    try:
        required_sections = [
            'research design',
            'data collection',
            'analysis framework',
            'ethical considerations',
        ]
        h2_lower = [h.lower() for h in headings_h2]
        sections_found = 0
        for req in required_sections:
            if any(req in h for h in h2_lower):
                sections_found += 1
                print(f"  PASS: Heading 2 '{req}' found")
            else:
                print(f"  FAIL: Heading 2 '{req}' not found. H2s present: {headings_h2}")

        section_score = (sections_found / 4.0) * 0.25
        if sections_found == 4:
            print(f"PASS: Component 2 — All 4 Heading 2 sections found ({section_score} pts)")
        else:
            print(f"PARTIAL: Component 2 — {sections_found}/4 sections found ({section_score} pts)")
        total_score += section_score
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: At least 2 footnotes (0.20 points)
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        footnote_refs = doc.element.body.findall('.//w:footnoteReference', ns)
        num_footnotes = len(footnote_refs)

        if num_footnotes >= 2:
            print(f"PASS: Component 3 — {num_footnotes} footnotes found (>=2 required) (0.20 pts)")
            total_score += 0.20
        elif num_footnotes == 1:
            print(f"PARTIAL: Component 3 — Only 1 footnote found (need >=2) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No footnotes found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Bibliography section with 3+ references (0.20 points)
    try:
        bib_section_found = False
        bib_entries = []
        for p in doc.paragraphs:
            if p.style and 'Heading' in p.style.name and 'bibliograph' in p.text.lower():
                bib_section_found = True
                continue
            if p.style and 'Heading' in p.style.name and bib_section_found:
                # Another heading after bibliography — stop collecting
                break
            if bib_section_found and p.text.strip():
                bib_entries.append(p.text.strip())

        if not bib_section_found:
            # Also check for "References" as heading
            for p in doc.paragraphs:
                if p.style and 'Heading' in p.style.name and 'reference' in p.text.lower():
                    bib_section_found = True
                    continue
                if p.style and 'Heading' in p.style.name and bib_section_found:
                    break
                if bib_section_found and p.text.strip():
                    bib_entries.append(p.text.strip())

        if bib_section_found and len(bib_entries) >= 3:
            # Check if entries look like APA format (author, year pattern)
            apa_count = 0
            for entry in bib_entries:
                # APA: Author(s) (Year). Title...
                if '(' in entry and ')' in entry:
                    apa_count += 1
            if apa_count >= 3:
                print(f"PASS: Component 4 — Bibliography with {len(bib_entries)} APA-style entries (0.20 pts)")
                total_score += 0.20
            elif apa_count >= 1:
                print(f"PARTIAL: Component 4 — Bibliography found but only {apa_count}/3 entries "
                      f"look APA-formatted (0.10 pts)")
                total_score += 0.10
            else:
                print(f"PARTIAL: Component 4 — Bibliography found with {len(bib_entries)} entries "
                      f"but none appear APA-formatted (0.05 pts)")
                total_score += 0.05
        elif bib_section_found:
            entry_score = min(len(bib_entries), 3) / 3.0 * 0.15
            print(f"PARTIAL: Component 4 — Bibliography section found but only {len(bib_entries)} "
                  f"entries (need 3) ({entry_score:.2f} pts)")
            total_score += entry_score
        else:
            print(f"FAIL: Component 4 — No Bibliography or References section found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Table of Contents present at beginning (0.20 points)
    try:
        toc_found = False
        # Check method 1: A heading with "Table of Contents" or "Contents" near the beginning
        for p in doc.paragraphs[:5]:  # Look in first 5 paragraphs
            if p.text.strip().lower() in ('table of contents', 'contents'):
                toc_found = True
                break

        # Check method 2: Look for TOC field code in XML
        if not toc_found:
            ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for instr in doc.element.body.iter(f'{{{ns_w}}}instrText'):
                if instr.text and 'TOC' in instr.text:
                    toc_found = True
                    break

        if toc_found:
            print(f"PASS: Component 5 — Table of Contents found at beginning of document (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 5 — No Table of Contents found at beginning of document")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved LibreOffice state before verification
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
