"""
Initial Setup: Add a tracked change comment spanning two paragraphs
Task ID: writer_lec_093
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_093'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Quarterly Business Performance Report', level=1)

    # Paragraph 1
    doc.add_paragraph(
        'Greenfield Technologies Inc. has experienced significant growth during '
        'the third quarter of 2025, driven primarily by expansion into the '
        'Asia-Pacific market and the successful launch of our CloudSync platform. '
        'Revenue increased by 23% compared to the same period last year, reaching '
        '$187.4 million.'
    )

    # Paragraph 2
    doc.add_paragraph(
        'Our engineering division, led by VP of Engineering Sarah Chen, delivered '
        'three major product releases ahead of schedule. The infrastructure team '
        'reduced average deployment time from 45 minutes to under 8 minutes through '
        'their containerization initiative, resulting in a 40% improvement in '
        'developer productivity metrics.'
    )

    # Paragraph 3
    doc.add_paragraph(
        'Customer acquisition costs decreased by 15% quarter-over-quarter, while '
        'customer lifetime value increased by 22%. The marketing team, under the '
        'direction of Marcus Johnson, implemented a new attribution model that '
        'provided clearer insights into campaign effectiveness across digital and '
        'traditional channels.'
    )

    # Paragraph 4
    doc.add_paragraph(
        'Human resources reported a net headcount increase of 47 employees, bringing '
        'total staff to 1,284. Employee retention remained strong at 94.2%, well '
        'above the industry average of 87.5%. The new mentorship program launched in '
        'July has already paired 156 participants across 12 departments.'
    )

    # Paragraph 5
    doc.add_paragraph(
        'The finance team completed the migration to the new ERP system on schedule. '
        'Accounts receivable days outstanding improved from 52 to 38 days. The board '
        'approved a $12 million capital expenditure budget for facility upgrades at '
        'our Austin and Toronto offices, scheduled for completion by Q2 2026.'
    )

    # Paragraph 6 - Legal-sensitive content
    doc.add_paragraph(
        'The company acknowledges potential exposure related to pending litigation '
        'filed by DataVault Corp. regarding alleged patent infringement of their '
        'distributed storage technology. Outside counsel at Morrison & Whitfield LLP '
        'has advised that damages, if awarded, could range from $8.5 million to '
        '$34 million depending on the court\'s interpretation of prior licensing '
        'agreements executed between 2019 and 2022.'
    )

    # Paragraph 7 - Legal-sensitive content
    doc.add_paragraph(
        'Furthermore, the regulatory compliance review initiated by the Federal Trade '
        'Commission concerning our data handling practices in the consumer segment '
        'remains unresolved. The company has submitted all requested documentation '
        'and anticipates a preliminary determination by January 2026. Any adverse '
        'finding could result in mandatory process changes and potential penalties '
        'estimated between $2.1 million and $15.8 million.'
    )

    # Paragraph 8
    doc.add_paragraph(
        'Looking ahead to Q4, the sales pipeline shows $246 million in qualified '
        'opportunities, representing a 31% increase over the same period last year. '
        'Key deals with Meridian Healthcare Systems and Pacific Northwest Utilities '
        'are expected to close before year-end, contributing an estimated $18.7 million '
        'in annual recurring revenue.'
    )

    # Paragraph 9
    doc.add_paragraph(
        'The product roadmap for 2026 includes the launch of our AI-powered analytics '
        'module, codenamed Project Aurora, which has completed beta testing with '
        '28 enterprise clients. Early feedback indicates strong demand, with 89% of '
        'beta participants expressing intent to upgrade their existing subscriptions '
        'to include the new capabilities.'
    )

    # Paragraph 10
    doc.add_paragraph(
        'In conclusion, Greenfield Technologies continues to demonstrate robust '
        'operational performance and strategic positioning. The executive leadership '
        'team remains committed to delivering sustainable growth while maintaining '
        'our commitment to innovation, customer success, and responsible corporate '
        'governance practices.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
