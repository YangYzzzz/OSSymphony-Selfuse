"""
Initial Setup: Insert five footnotes in the literature review section
Task ID: writer_struct_066
Domain: libreoffice_writer
Creates: /home/user/Desktop/psychology_thesis.docx (NO footnotes - initial state)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_struct_066'
OUTPUT = f'{WORKDIR}/psychology_thesis.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    return para


def add_body_para(doc, text, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
    para = doc.add_paragraph(text)
    para.paragraph_format.alignment = alignment
    para.paragraph_format.first_line_indent = Inches(0.5)
    para.paragraph_format.space_after = Pt(6)
    return para


def create_initial():
    # Create desktop directory if needed
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Page setup - standard thesis margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ================================================================
    # TITLE PAGE (Page 1)
    # ================================================================
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run("The Psychological Mechanisms of Cognitive Resilience:\nA Theoretical and Empirical Investigation")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()

    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = sub_para.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy")
    run2.font.size = Pt(12)

    doc.add_paragraph()
    author_para = doc.add_paragraph()
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = author_para.add_run("Emily Hartwell\nDepartment of Psychology\nUniversity of Westfield\n2025")
    run3.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # ABSTRACT (Page 2)
    # ================================================================
    add_heading(doc, "Abstract", level=1)

    abstract_text = (
        "This thesis investigates the psychological mechanisms underlying cognitive resilience "
        "in adults aged 25-65 who have experienced significant adverse life events. Drawing on "
        "a mixed-methods approach combining longitudinal survey data (N=847) with qualitative "
        "interviews (n=42), the study identifies four primary protective factors: adaptive "
        "coping strategies, robust social support networks, growth-oriented cognitive appraisal, "
        "and neuroplasticity-enabling behavioral routines. The findings suggest that resilience "
        "is not a fixed trait but a dynamic process that can be cultivated through targeted "
        "psychological interventions. Implications for clinical practice and community-based "
        "programs are discussed, along with directions for future research."
    )
    add_body_para(doc, abstract_text)

    keywords_para = doc.add_paragraph()
    run_kw = keywords_para.add_run("Keywords: ")
    run_kw.bold = True
    keywords_para.add_run("cognitive resilience, adverse life events, coping strategies, longitudinal study, psychological interventions")

    doc.add_page_break()

    # ================================================================
    # CHAPTER 1: INTRODUCTION (Page 3 starts here)
    # ================================================================
    add_heading(doc, "Chapter 1: Introduction", level=1)

    intro1 = (
        "Cognitive resilience — the capacity of the human mind to adapt, recover, and even "
        "flourish in the face of adversity — has emerged as one of the most consequential "
        "constructs in contemporary psychological science. Over the past three decades, "
        "researchers have moved beyond simplistic models of stress and vulnerability toward "
        "more nuanced frameworks that account for protective factors, individual differences, "
        "and contextual determinants of psychological well-being."
    )
    add_body_para(doc, intro1)

    intro2 = (
        "The impetus for this investigation arose from observations in clinical practice "
        "where individuals confronting seemingly insurmountable challenges — chronic illness, "
        "bereavement, financial collapse, relational dissolution — demonstrated markedly "
        "divergent trajectories of recovery. While some individuals exhibited prolonged "
        "psychological distress and functional impairment, others displayed remarkable "
        "adaptive capacity, recovering their baseline functioning within months and, in "
        "some cases, reporting post-traumatic growth."
    )
    add_body_para(doc, intro2)

    intro3 = (
        "The present thesis seeks to address a critical gap in the existing literature: "
        "the absence of an integrative theoretical model that accounts for both dispositional "
        "and situational determinants of resilience across the adult lifespan. By synthesizing "
        "empirical findings from cognitive neuroscience, developmental psychology, and clinical "
        "intervention research, this thesis proposes a multi-level framework of cognitive "
        "resilience that encompasses biological, psychological, and social dimensions."
    )
    add_body_para(doc, intro3)

    # ================================================================
    # CHAPTER 2: LITERATURE REVIEW (Pages 3-5)
    # ================================================================
    add_heading(doc, "Chapter 2: Literature Review", level=1)

    add_heading(doc, "2.1 Historical Perspectives on Resilience Research", level=2)

    lit_hist1 = (
        "Early scholarly interest in resilience emerged from developmental psychology, "
        "particularly from studies of children raised in adverse environments. Researchers "
        "in the 1970s and 1980s observed that a substantial proportion of children exposed "
        "to poverty, abuse, and family dysfunction did not develop the expected psychiatric "
        "sequelae — a phenomenon that challenged prevailing deficit-oriented models of "
        "psychological development according to prior research. These foundational observations "
        "catalyzed decades of empirical inquiry into the nature and determinants of adaptive "
        "functioning under adversity."
    )
    add_body_para(doc, lit_hist1)

    lit_hist2 = (
        "Subsequent theoretical elaborations shifted the conceptualization of resilience from "
        "a static trait to a dynamic process. Werner and colleagues' landmark longitudinal "
        "study of children born into high-risk environments on the island of Kauai provided "
        "compelling evidence that protective factors — including temperamental characteristics, "
        "familial cohesion, and community resources — could buffer the deleterious effects of "
        "chronic adversity. These theoretical frameworks were later validated in clinical settings "
        "as researchers extended this work to adult populations grappling with trauma, chronic "
        "illness, and occupational stress."
    )
    add_body_para(doc, lit_hist2)

    lit_hist3 = (
        "The conceptual expansion of resilience research into adulthood necessitated "
        "methodological refinements capable of capturing the dynamic interplay between "
        "biological, psychological, and social processes over extended time periods. "
        "Cross-sectional designs, while informative, proved insufficient to capture "
        "trajectories of adaptation. Longitudinal studies, though methodologically demanding, "
        "became the gold standard for resilience research. Integrating findings across "
        "diverse study populations and methodologies, meta-analytic evidence suggests that "
        "resilience is a multidimensional construct with moderate heritability and substantial "
        "environmental modifiability, particularly during sensitive developmental periods."
    )
    add_body_para(doc, lit_hist3)

    add_heading(doc, "2.2 Cognitive and Neurobiological Correlates", level=2)

    lit_cog1 = (
        "Advances in cognitive neuroscience have substantially enriched our understanding "
        "of the neural substrates underlying resilient responding to adversity. Neuroimaging "
        "studies have consistently identified prefrontal cortex activity as a key correlate "
        "of adaptive emotion regulation, with resilient individuals demonstrating greater "
        "prefrontal-amygdala connectivity during emotional challenge tasks. Executive function "
        "capacities — including working memory, cognitive flexibility, and inhibitory control — "
        "have emerged as central mediators of the relationship between adversity exposure and "
        "psychological outcome. Across a diverse array of samples spanning different cultural "
        "contexts and socioeconomic strata, longitudinal data confirms that executive function "
        "training in mid-life is associated with significantly reduced rates of anxiety and "
        "depressive symptomatology over five-year follow-up periods."
    )
    add_body_para(doc, lit_cog1)

    lit_cog2 = (
        "The role of neuroplasticity in sustaining resilient cognitive functioning across "
        "the adult lifespan has attracted considerable scientific attention. Animal and human "
        "studies converge on the finding that neuroplasticity-enhancing behaviors — regular "
        "aerobic exercise, cognitive novelty-seeking, mindfulness-based practices — can "
        "attenuate age-related cognitive decline and bolster stress resilience. Structural "
        "MRI studies have documented greater hippocampal volume in individuals with high "
        "resilience scores, consistent with the hypothesis that neurogenic processes may "
        "mediate the protective effects of positive affect and social engagement."
    )
    add_body_para(doc, lit_cog2)

    add_heading(doc, "2.3 Social and Contextual Determinants", level=2)

    lit_soc1 = (
        "The social embeddedness of resilience represents one of the most robust and "
        "replicated findings in the field. Perceived social support — encompassing emotional "
        "support, informational support, and tangible assistance — has been identified as "
        "a primary protective factor across the lifespan. Importantly, the perceived quality "
        "of social relationships appears to be more predictive of resilient outcomes than "
        "the objective quantity of social contacts, suggesting that subjective appraisals "
        "of social resources mediate the protective effect of social networks."
    )
    add_body_para(doc, lit_soc1)

    lit_soc2 = (
        "Cultural context constitutes a frequently underexamined moderator of resilience "
        "processes. Cross-cultural comparative studies have documented significant variation "
        "in the relative salience of individual versus collective coping orientations, "
        "with collectivist cultures emphasizing interdependent resilience strategies that "
        "leverage family and community resources more heavily than individualistic cultures. "
        "The universality of resilience mechanisms — specifically, whether core protective "
        "factors operate similarly across cultural contexts — has been contested by recent "
        "replication attempts, which have highlighted the importance of culturally-specific "
        "operationalizations of constructs such as coping efficacy and post-traumatic growth."
    )
    add_body_para(doc, lit_soc2)

    lit_soc3 = (
        "Socioeconomic factors intersect with individual and social resources to shape "
        "resilience trajectories in complex ways. Economic adversity not only constitutes "
        "a major stressor in its own right but also depletes the cognitive, emotional, and "
        "material resources available for adaptive coping. Poverty-related chronic stress "
        "can compromise the neurobiological substrates of resilience — including "
        "hypothalamic-pituitary-adrenal axis regulation and prefrontal cortex functioning — "
        "creating a vicious cycle of vulnerability and adversity."
    )
    add_body_para(doc, lit_soc3)

    doc.add_page_break()

    # ================================================================
    # CHAPTER 3: METHODOLOGY (Pages 5-7)
    # ================================================================
    add_heading(doc, "Chapter 3: Methodology", level=1)

    add_heading(doc, "3.1 Research Design", level=2)

    meth1 = (
        "This investigation employed a sequential explanatory mixed-methods design, integrating "
        "a longitudinal quantitative survey component with a qualitative interview component. "
        "The mixed-methods approach was selected to provide both the statistical power necessary "
        "for examining population-level associations between resilience factors and outcomes, "
        "and the interpretive depth needed to understand the lived experience of cognitive "
        "resilience in diverse adult populations."
    )
    add_body_para(doc, meth1)

    meth2 = (
        "The quantitative phase involved three waves of data collection over a 24-month period "
        "(T1: baseline, T2: 12 months, T3: 24 months). The qualitative phase comprised "
        "in-depth semi-structured interviews conducted with a theoretically-sampled subset of "
        "survey participants following completion of the T3 survey. Integration of findings "
        "occurred at the interpretation stage, with qualitative themes used to explain, "
        "contextualize, and extend the quantitative results."
    )
    add_body_para(doc, meth2)

    add_heading(doc, "3.2 Participants", level=2)

    part1 = (
        "The quantitative sample comprised 847 adults aged 25-65 (M=43.7, SD=11.2) recruited "
        "through a stratified community sampling procedure across three metropolitan regions. "
        "Inclusion criteria required: (1) age between 25 and 65 years at baseline; (2) "
        "self-reported experience of at least one significant adverse life event in the "
        "preceding five years, as assessed by the Life Events Checklist; and (3) sufficient "
        "proficiency in English to complete the survey instruments."
    )
    add_body_para(doc, part1)

    part2 = (
        "The qualitative sub-sample (n=42) was selected to maximize variation across key "
        "demographic dimensions (age, gender, ethnicity, socioeconomic status) and resilience "
        "trajectories (high resilience, moderate resilience, low resilience). This purposive "
        "sampling strategy was designed to illuminate the mechanisms through which resilience "
        "resources operate across diverse life circumstances and adversity types."
    )
    add_body_para(doc, part2)

    add_heading(doc, "3.3 Measures", level=2)

    meas1 = (
        "Cognitive resilience was operationalized using the Connor-Davidson Resilience Scale "
        "(CD-RISC; 25 items, alpha=.89), supplemented by the Brief Resilience Scale (BRS; "
        "6 items, alpha=.87) to capture both trait-like and process-oriented dimensions of "
        "adaptive functioning. Adverse life experiences were assessed using the Life Events "
        "Checklist for DSM-5 (LEC-5), with additional modules assessing cumulative adversity "
        "exposure across the lifespan."
    )
    add_body_para(doc, meas1)

    doc.add_page_break()

    # ================================================================
    # CHAPTER 4: RESULTS (Pages 7-10)
    # ================================================================
    add_heading(doc, "Chapter 4: Results", level=1)

    add_heading(doc, "4.1 Descriptive Statistics and Sample Characteristics", level=2)

    res1 = (
        "At baseline (T1), the sample demonstrated a mean resilience score of 68.4 (SD=12.7) "
        "on the CD-RISC, with scores ranging from 31 to 100. This distribution approximated "
        "normality (skewness=0.14, kurtosis=0.23) and was broadly comparable to normative "
        "data reported for community-dwelling adults in previous studies. Women comprised "
        "54.3% of the sample, and the sample was broadly representative of the metropolitan "
        "population across measures of age, education, and occupational status."
    )
    add_body_para(doc, res1)

    res2 = (
        "Among the 847 participants, the most frequently endorsed adverse life events at "
        "baseline were: loss of a close family member (63.2%), serious personal illness or "
        "injury (41.7%), major financial crisis (38.5%), relationship dissolution (35.9%), "
        "and job loss or career disruption (31.4%). A substantial proportion of participants "
        "had experienced multiple adverse events (M=2.84, SD=1.37), reflecting the cumulative "
        "adversity exposure characterizing many adult lives."
    )
    add_body_para(doc, res2)

    add_heading(doc, "4.2 Longitudinal Trajectories of Resilience", level=2)

    res3 = (
        "Latent growth curve analysis identified three distinct trajectories of resilience "
        "over the 24-month follow-up period. The largest group (52.7%, n=446) demonstrated "
        "stable high resilience across all three time points, with minimal variation from "
        "baseline levels. A second group (31.4%, n=266) exhibited an initial resilience "
        "decrement at T2, followed by recovery to near-baseline levels at T3 — a trajectory "
        "consistent with theoretical models of dynamic resilience. The smallest group (15.9%, "
        "n=135) displayed persistent low resilience scores across the observation period, "
        "with modest but statistically non-significant improvement over time."
    )
    add_body_para(doc, res3)

    res4 = (
        "Predictors of trajectory group membership were examined using multinomial logistic "
        "regression. Higher baseline executive function scores (OR=2.34, 95% CI [1.87, 2.93], "
        "p<.001), greater perceived social support (OR=1.89, 95% CI [1.52, 2.35], p<.001), "
        "and more frequent engagement in neuroplasticity-enhancing behaviors (OR=1.67, 95% CI "
        "[1.34, 2.08], p<.001) were each independently associated with membership in the "
        "stable high resilience trajectory, controlling for demographic covariates."
    )
    add_body_para(doc, res4)

    doc.add_page_break()

    # ================================================================
    # CHAPTER 5: DISCUSSION (Pages 10-12)
    # ================================================================
    add_heading(doc, "Chapter 5: Discussion", level=1)

    add_heading(doc, "5.1 Synthesis of Findings", level=2)

    disc1 = (
        "The present findings converge with and extend prior literature in several important "
        "respects. First, the identification of three distinct resilience trajectories "
        "empirically validates theoretical models proposing that resilience is not a unitary "
        "construct but a heterogeneous phenomenon with multiple developmental pathways. "
        "The stable high resilience trajectory, observed in over half the sample, is consistent "
        "with theoretical formulations emphasizing resilience as a maintained adaptive capacity "
        "rather than a recovery process per se."
    )
    add_body_para(doc, disc1)

    disc2 = (
        "Second, the identification of executive function, social support, and "
        "neuroplasticity-enabling behaviors as independent predictors of trajectory membership "
        "provides empirical support for multi-level models of resilience. These findings "
        "suggest that resilience is multiply determined — simultaneously shaped by "
        "neurobiological capacities, interpersonal resources, and behavioral practices — "
        "and that interventions targeting any of these dimensions may enhance adaptive "
        "functioning under adversity."
    )
    add_body_para(doc, disc2)

    add_heading(doc, "5.2 Clinical and Policy Implications", level=2)

    disc3 = (
        "These findings carry significant implications for the design of resilience-enhancing "
        "interventions across clinical and community settings. Cognitive training programs "
        "targeting executive function — including working memory training, attentional control "
        "interventions, and cognitive-behavioral stress management — may yield particularly "
        "robust effects for individuals at risk of persistent low resilience, for whom "
        "neurobiological vulnerabilities represent a primary point of intervention."
    )
    add_body_para(doc, disc3)

    disc4 = (
        "At the policy level, the findings underscore the importance of structural supports "
        "for resilience in the face of socioeconomic adversity. Programs designed to reduce "
        "financial insecurity, strengthen community social infrastructure, and expand access "
        "to mental health services may generate cascading benefits for population-level "
        "resilience by attenuating the resource-depleting effects of chronic economic stress."
    )
    add_body_para(doc, disc4)

    doc.add_page_break()

    # ================================================================
    # CHAPTER 6: CONCLUSION (Pages 12-13)
    # ================================================================
    add_heading(doc, "Chapter 6: Conclusion", level=1)

    conc1 = (
        "This thesis has presented an empirical investigation of cognitive resilience in "
        "adulthood, drawing on longitudinal survey data and qualitative interviews to "
        "identify protective factors and adaptive trajectories. The central contribution "
        "of this work is an empirically-derived, multi-level model of resilience that "
        "integrates neurobiological, psychological, and social determinants within a "
        "dynamic developmental framework."
    )
    add_body_para(doc, conc1)

    conc2 = (
        "Cognitive resilience is best understood not as a fixed dispositional attribute "
        "but as a dynamic capacity shaped by ongoing transactions between individuals and "
        "their biological, social, and cultural environments. The implications of this "
        "perspective for intervention design are substantial: rather than targeting a single "
        "domain of protective functioning, resilience-enhancing programs may achieve the "
        "greatest impact by adopting multi-component, ecologically-situated approaches "
        "that address the full complexity of adaptive functioning in context."
    )
    add_body_para(doc, conc2)

    conc3 = (
        "Future research should prioritize the development of validated, culturally-adapted "
        "resilience measures capable of capturing the diversity of adaptive strategies across "
        "different populations. Additionally, mechanistic studies employing neuroimaging and "
        "behavioral methods are needed to elucidate the causal pathways through which "
        "protective factors exert their beneficial effects — and to identify optimal timing "
        "windows for resilience-enhancing interventions across the adult lifespan."
    )
    add_body_para(doc, conc3)

    doc.add_page_break()

    # ================================================================
    # REFERENCES (Pages 13-15)
    # ================================================================
    add_heading(doc, "References", level=1)

    refs = [
        "Bonanno, G. A. (2004). Loss, trauma, and human resilience: Have we underestimated the human capacity to thrive after extremely aversive events? American Psychologist, 59(1), 20-28.",
        "Cicchetti, D., & Garmezy, N. (1993). Prospects and promises in the study of resilience. Development and Psychopathology, 5(4), 497-502.",
        "Connor, K. M., & Davidson, J. R. T. (2003). Development of a new resilience scale: The Connor-Davidson Resilience Scale (CD-RISC). Depression and Anxiety, 18(2), 76-82.",
        "Fletcher, D., & Sarkar, M. (2013). Psychological resilience: A review and critique of definitions, concepts, and theory. European Psychologist, 18(1), 12-23.",
        "Garmezy, N. (1991). Resilience in children's adaptation to negative life events and stressed environments. Pediatric Annals, 20(9), 459-466.",
        "Masten, A. S. (2001). Ordinary magic: Resilience processes in development. American Psychologist, 56(3), 227-238.",
        "Rutter, M. (2006). Implications of resilience concepts for scientific understanding. Annals of the New York Academy of Sciences, 1094(1), 1-12.",
        "Southwick, S. M., & Charney, D. S. (2012). The science of resilience: Implications for the prevention and treatment of depression. Science, 338(6103), 79-82.",
        "Tugade, M. M., & Fredrickson, B. L. (2004). Resilient individuals use positive emotions to bounce back from negative emotional experiences. Journal of Personality and Social Psychology, 86(2), 320-333.",
        "Werner, E. E., & Smith, R. S. (1992). Overcoming the odds: High-risk children from birth to adulthood. Cornell University Press.",
        "Yehuda, R., & Flory, J. D. (2007). Differentiating biological correlates of risk, PTSD, and resilience following trauma exposure. Journal of Traumatic Stress, 20(4), 435-447.",
        "Zautra, A. J., Hall, J. S., & Murray, K. E. (2010). Resilience: A new definition of health for people and communities. In J. W. Reich, A. J. Zautra, & J. S. Hall (Eds.), Handbook of adult resilience (pp. 3-34). Guilford Press.",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref, style='List Paragraph')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)

    # Save the file
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
