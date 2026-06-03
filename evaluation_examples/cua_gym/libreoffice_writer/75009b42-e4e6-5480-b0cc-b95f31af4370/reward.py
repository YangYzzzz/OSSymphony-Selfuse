"""
Reward Script: Corporate document template with custom styles
Task ID: writer_biz_060
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.15): Three Corp styles exist
  - Component 2 (0.15): Corp Body style properties correct
  - Component 3 (0.15): Corp Heading 1 style properties correct
  - Component 4 (0.15): Corp Heading 2 style properties correct
  - Component 5 (0.15): Body paragraphs assigned Corp Body style
  - Component 6 (0.125): Heading 1 paragraphs assigned Corp Heading 1 style
  - Component 7 (0.125): Heading 2 paragraphs assigned Corp Heading 2 style
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_060'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect style names
    style_names = set()
    for style in doc.styles:
        if style.type is not None and style.type.name == 'PARAGRAPH':
            style_names.add(style.name)

    # Component 1: Three Corp styles exist (0.15 points)
    try:
        required_styles = ['Corp Body', 'Corp Heading 1', 'Corp Heading 2']
        found = [s for s in required_styles if s in style_names]
        if len(found) == 3:
            print(f"PASS: Component 1 -- All 3 Corp styles exist: {found} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 -- Only found {len(found)}/3 Corp styles: {found}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Corp Body style properties (0.15 points)
    # Expected: Calibri 11pt, justified, 1.15 line spacing, 6pt after
    try:
        if 'Corp Body' in style_names:
            s = doc.styles['Corp Body']
            f = s.font
            pf = s.paragraph_format
            sub_checks = 0
            sub_total = 4

            # Font name
            if f.name and f.name.lower() == 'calibri':
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Body font.name={f.name}, expected Calibri")

            # Font size 11pt
            if f.size and abs(f.size.pt - 11.0) < 0.5:
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Body font.size={f.size.pt if f.size else None}, expected 11.0")

            # Justified alignment
            if pf.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Body alignment={pf.alignment}, expected JUSTIFY")

            # Line spacing 1.15 AND space_after 6pt (combined as one sub-check)
            if (pf.line_spacing is not None and abs(float(pf.line_spacing) - 1.15) < 0.05
                    and pf.space_after is not None and abs(pf.space_after.pt - 6.0) < 0.5):
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Body line_spacing={pf.line_spacing}, space_after={pf.space_after.pt if pf.space_after else None}")

            if sub_checks > 0:
                pts = 0.15 * (sub_checks / sub_total)
                if sub_checks == sub_total:
                    print(f"PASS: Component 2 -- Corp Body style properties all correct (0.15 pts)")
                else:
                    print(f"PARTIAL: Component 2 -- Corp Body {sub_checks}/{sub_total} sub-checks ({pts:.3f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 2 -- Corp Body 0/{sub_total} sub-checks passed")
        else:
            print(f"FAIL: Component 2 -- Corp Body style does not exist")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Corp Heading 1 style properties (0.15 points)
    # Expected: Calibri 16pt bold, #003366, 18pt before, 6pt after
    try:
        if 'Corp Heading 1' in style_names:
            s = doc.styles['Corp Heading 1']
            f = s.font
            pf = s.paragraph_format
            sub_checks = 0
            sub_total = 4

            # Font name + size
            if f.name and f.name.lower() == 'calibri' and f.size and abs(f.size.pt - 16.0) < 0.5:
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 1 font={f.name} size={f.size.pt if f.size else None}")

            # Bold
            if f.bold is True:
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 1 bold={f.bold}")

            # Color #003366
            if f.color and f.color.rgb and str(f.color.rgb).upper() == '003366':
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 1 color={f.color.rgb if f.color else None}")

            # Spacing: 18pt before, 6pt after
            if (pf.space_before is not None and abs(pf.space_before.pt - 18.0) < 1.0
                    and pf.space_after is not None and abs(pf.space_after.pt - 6.0) < 1.0):
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 1 space_before={pf.space_before.pt if pf.space_before else None}, space_after={pf.space_after.pt if pf.space_after else None}")

            if sub_checks > 0:
                pts = 0.15 * (sub_checks / sub_total)
                if sub_checks == sub_total:
                    print(f"PASS: Component 3 -- Corp Heading 1 style properties all correct (0.15 pts)")
                else:
                    print(f"PARTIAL: Component 3 -- Corp Heading 1 {sub_checks}/{sub_total} sub-checks ({pts:.3f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 3 -- Corp Heading 1 0/{sub_total} sub-checks passed")
        else:
            print(f"FAIL: Component 3 -- Corp Heading 1 style does not exist")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Corp Heading 2 style properties (0.15 points)
    # Expected: Calibri 13pt bold, #336699, 12pt before, 4pt after
    try:
        if 'Corp Heading 2' in style_names:
            s = doc.styles['Corp Heading 2']
            f = s.font
            pf = s.paragraph_format
            sub_checks = 0
            sub_total = 4

            # Font name + size
            if f.name and f.name.lower() == 'calibri' and f.size and abs(f.size.pt - 13.0) < 0.5:
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 2 font={f.name} size={f.size.pt if f.size else None}")

            # Bold
            if f.bold is True:
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 2 bold={f.bold}")

            # Color #336699
            if f.color and f.color.rgb and str(f.color.rgb).upper() == '336699':
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 2 color={f.color.rgb if f.color else None}")

            # Spacing: 12pt before, 4pt after
            if (pf.space_before is not None and abs(pf.space_before.pt - 12.0) < 1.0
                    and pf.space_after is not None and abs(pf.space_after.pt - 4.0) < 1.0):
                sub_checks += 1
            else:
                print(f"  DETAIL: Corp Heading 2 space_before={pf.space_before.pt if pf.space_before else None}, space_after={pf.space_after.pt if pf.space_after else None}")

            if sub_checks > 0:
                pts = 0.15 * (sub_checks / sub_total)
                if sub_checks == sub_total:
                    print(f"PASS: Component 4 -- Corp Heading 2 style properties all correct (0.15 pts)")
                else:
                    print(f"PARTIAL: Component 4 -- Corp Heading 2 {sub_checks}/{sub_total} sub-checks ({pts:.3f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 4 -- Corp Heading 2 0/{sub_total} sub-checks passed")
        else:
            print(f"FAIL: Component 4 -- Corp Heading 2 style does not exist")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: Body paragraphs assigned Corp Body (0.15 points)
    # In the initial doc, body paragraphs use 'Normal'. In golden, they should use 'Corp Body'.
    try:
        body_paras = [p for p in doc.paragraphs if p.text.strip()]
        # Identify paragraphs that should be body (not headings)
        # A body paragraph is one that is NOT a heading-level paragraph
        body_count = 0
        corp_body_count = 0
        for p in body_paras:
            sname = p.style.name
            # If it's not any heading style, it should be Corp Body
            if not any(h in sname.lower() for h in ['heading', 'title']):
                body_count += 1
                if sname == 'Corp Body':
                    corp_body_count += 1

        if body_count > 0:
            ratio = corp_body_count / body_count
            if ratio >= 0.9:
                print(f"PASS: Component 5 -- {corp_body_count}/{body_count} body paragraphs use Corp Body (0.15 pts)")
                total_score += 0.15
            elif ratio > 0:
                pts = 0.15 * ratio
                print(f"PARTIAL: Component 5 -- {corp_body_count}/{body_count} body paragraphs use Corp Body ({pts:.3f} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 5 -- 0/{body_count} body paragraphs use Corp Body")
        else:
            print(f"FAIL: Component 5 -- No body paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    # Component 6: Heading 1 paragraphs assigned Corp Heading 1 (0.125 points)
    # Initial uses 'Heading 1', golden should use 'Corp Heading 1'
    try:
        h1_paras = [p for p in doc.paragraphs if p.text.strip()]
        # Count paragraphs with Corp Heading 1 style
        corp_h1_count = sum(1 for p in h1_paras if p.style.name == 'Corp Heading 1')
        # There should be at least 1 Corp Heading 1 paragraph (we saw 4 in golden)
        if corp_h1_count >= 2:
            print(f"PASS: Component 6 -- {corp_h1_count} paragraphs use Corp Heading 1 (0.125 pts)")
            total_score += 0.125
        elif corp_h1_count >= 1:
            print(f"PARTIAL: Component 6 -- Only {corp_h1_count} paragraph uses Corp Heading 1 (0.0625 pts)")
            total_score += 0.0625
        else:
            print(f"FAIL: Component 6 -- No paragraphs use Corp Heading 1 style")
    except Exception as e:
        print(f"ERROR: Component 6 -- {e}")

    # Component 7: Heading 2 paragraphs assigned Corp Heading 2 (0.125 points)
    try:
        h2_paras = [p for p in doc.paragraphs if p.text.strip()]
        corp_h2_count = sum(1 for p in h2_paras if p.style.name == 'Corp Heading 2')
        if corp_h2_count >= 2:
            print(f"PASS: Component 7 -- {corp_h2_count} paragraphs use Corp Heading 2 (0.125 pts)")
            total_score += 0.125
        elif corp_h2_count >= 1:
            print(f"PARTIAL: Component 7 -- Only {corp_h2_count} paragraph uses Corp Heading 2 (0.0625 pts)")
            total_score += 0.0625
        else:
            print(f"FAIL: Component 7 -- No paragraphs use Corp Heading 2 style")
    except Exception as e:
        print(f"ERROR: Component 7 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
