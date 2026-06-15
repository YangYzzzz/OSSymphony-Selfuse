"""
Initial Setup: Corporate document with default styles
Task ID: writer_biz_060
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_060'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Use default styles only - no custom styles
    # The document is a quarterly business report with default formatting

    # --- Main Heading 1 ---
    doc.add_heading('Quarterly Business Performance Report', level=1)

    # --- Body text ---
    doc.add_paragraph(
        'This report provides a comprehensive overview of our company\'s performance '
        'during Q4 2025. The analysis covers revenue trends, operational efficiency, '
        'market expansion efforts, and strategic initiatives undertaken across all '
        'business divisions.'
    )

    doc.add_paragraph(
        'All figures presented in this document have been reviewed by the internal '
        'audit team and reconciled against our enterprise resource planning system. '
        'The data reflects actual results as of December 31, 2025.'
    )

    # --- Heading 1 ---
    doc.add_heading('Financial Overview', level=1)

    # --- Body text ---
    doc.add_paragraph(
        'Total consolidated revenue for Q4 2025 reached $127.4 million, representing '
        'a 14.2% increase over the same period last year. This growth was primarily '
        'driven by strong performance in our enterprise software division and '
        'expanding partnerships in the Asia-Pacific region.'
    )

    # --- Heading 2 ---
    doc.add_heading('Revenue Breakdown by Division', level=2)

    doc.add_paragraph(
        'The Enterprise Solutions division contributed $68.3 million, accounting for '
        '53.6% of total revenue. The Cloud Services segment generated $34.1 million, '
        'up 22% year-over-year, while Professional Services brought in $25.0 million.'
    )

    doc.add_paragraph(
        'Notably, recurring subscription revenue now represents 71% of our total '
        'income, up from 62% in Q4 2024. This shift toward predictable revenue '
        'streams strengthens our financial stability and improves forecasting accuracy.'
    )

    # --- Heading 2 ---
    doc.add_heading('Cost Analysis and Margins', level=2)

    doc.add_paragraph(
        'Gross margin improved to 68.4%, compared to 65.1% in the prior year quarter. '
        'This improvement reflects ongoing optimization of our cloud infrastructure '
        'costs and favorable changes in our revenue mix toward higher-margin products.'
    )

    doc.add_paragraph(
        'Operating expenses totaled $62.8 million, with research and development '
        'accounting for $28.5 million (45.4% of OpEx). Sales and marketing expenses '
        'were $22.1 million, and general administrative costs were $12.2 million.'
    )

    # --- Heading 1 ---
    doc.add_heading('Operational Highlights', level=1)

    doc.add_paragraph(
        'Several key operational milestones were achieved during the quarter, '
        'including the successful launch of our next-generation analytics platform '
        'and the completion of our data center expansion in Frankfurt, Germany.'
    )

    # --- Heading 2 ---
    doc.add_heading('Product Development', level=2)

    doc.add_paragraph(
        'Our engineering teams shipped 47 feature releases during Q4, including '
        'the highly anticipated real-time collaboration module. Customer adoption '
        'of the new features exceeded projections, with 38% of enterprise clients '
        'activating the collaboration tools within the first 30 days.'
    )

    doc.add_paragraph(
        'The AI-powered document analysis feature entered beta testing with '
        '150 selected customers. Early feedback indicates a 40% reduction in '
        'document processing time and a 95% accuracy rate in data extraction.'
    )

    # --- Heading 2 ---
    doc.add_heading('Customer Success Metrics', level=2)

    doc.add_paragraph(
        'Net Promoter Score improved to 72, up from 68 in Q3. Customer retention '
        'rate remained strong at 94.3%, and the average contract value for new '
        'enterprise deals increased by 18% to $245,000 annually.'
    )

    doc.add_paragraph(
        'Our support team resolved 96.1% of critical tickets within the 4-hour SLA, '
        'and overall customer satisfaction with support interactions reached 4.7 out '
        'of 5.0 based on post-interaction surveys.'
    )

    # --- Heading 1 ---
    doc.add_heading('Strategic Outlook', level=1)

    doc.add_paragraph(
        'Looking ahead to Q1 2026, we anticipate continued momentum in our enterprise '
        'segment, driven by pipeline expansion in healthcare and financial services '
        'verticals. Capital expenditure is projected at $15.2 million for data center '
        'upgrades and new AI infrastructure.'
    )

    # --- Heading 2 ---
    doc.add_heading('Market Expansion Plans', level=2)

    doc.add_paragraph(
        'We plan to establish regional offices in Singapore and Toronto during the '
        'first half of 2026, targeting the Southeast Asian and Canadian markets '
        'respectively. Initial investment for these expansions is budgeted at '
        '$8.5 million, with expected breakeven within 18 months.'
    )

    doc.add_paragraph(
        'Strategic partnerships with three major system integrators are in advanced '
        'negotiations, which would extend our reach into government and defense '
        'sectors across five additional countries.'
    )

    # --- Heading 2 ---
    doc.add_heading('Technology Roadmap', level=2)

    doc.add_paragraph(
        'Key technology investments for 2026 include the migration of our core '
        'platform to a microservices architecture, expected to reduce deployment '
        'cycles from weeks to hours. Additionally, we are investing $12 million in '
        'our proprietary large language model for industry-specific applications.'
    )

    doc.add_paragraph(
        'Security enhancements remain a top priority, with planned implementation '
        'of zero-trust architecture across all customer-facing services by Q3 2026. '
        'This initiative has been allocated a dedicated budget of $4.8 million.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
