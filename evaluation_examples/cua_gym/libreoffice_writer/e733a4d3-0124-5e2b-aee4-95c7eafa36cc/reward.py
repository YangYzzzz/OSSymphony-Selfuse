"""
Reward Script: Complete mail merge setup for event invitation letter
Task ID: osworld_writer_mail_merge_003
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): No placeholder markers remain (<<Name>>, <<EventDate>>, <<Venue>> all replaced)
  Component 2 (0.3): All 3 recipient names correctly merged (Alice Hoffman, Benjamin Clarke, Catherine Nguyen)
  Component 3 (0.2): Correct dates and venues for each of the 3 records
  Component 4 (0.1): Document has 2 page breaks separating the 3 records
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_mail_merge_003'

# Ground truth from context: attendees.csv records 1-3
EXPECTED_RECORDS = [
    {
        'name': 'Alice Hoffman',
        'date': 'April 12, 2025',
        'venue': 'Grand Ballroom, Ritz Hotel',
    },
    {
        'name': 'Benjamin Clarke',
        'date': 'April 12, 2025',
        'venue': 'Grand Ballroom, Ritz Hotel',
    },
    {
        'name': 'Catherine Nguyen',
        'date': 'April 13, 2025',
        'venue': 'Skyline Terrace, Blue Tower',
    },
]

PLACEHOLDER_MARKERS = ['<<Name>>', '<<EventDate>>', '<<Venue>>']


def verify_task(file_path):
    """
    Verify mail merge task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph text for searching
    all_text = '\n'.join(para.text for para in doc.paragraphs)

    # Component 1: No placeholder markers remain (0.4 points)
    # The task requires replacing <<Name>>, <<EventDate>>, <<Venue>> with real values.
    # Initial file has these markers; golden file should have them all replaced.
    try:
        placeholders_remaining = [p for p in PLACEHOLDER_MARKERS if p in all_text]
        if not placeholders_remaining:
            print("PASS: Component 1 — No placeholder markers remain in the document (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Placeholder markers still present: {placeholders_remaining}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 3 recipient names correctly merged (0.3 points)
    # Each name must appear in a 'Dear <Name>,' salutation line.
    try:
        names_found = []
        for record in EXPECTED_RECORDS:
            expected_salutation = f"Dear {record['name']},"
            found = any(expected_salutation in para.text for para in doc.paragraphs)
            if found:
                names_found.append(record['name'])

        if len(names_found) == 3:
            print(f"PASS: Component 2 — All 3 recipient names correctly merged: {names_found} (0.3 pts)")
            total_score += 0.3
        elif len(names_found) > 0:
            print(f"PARTIAL-FAIL: Component 2 — Only {len(names_found)}/3 names found: {names_found}")
        else:
            print(f"FAIL: Component 2 — None of the expected recipient names found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Correct dates and venues for each record (0.2 points)
    # Each record's date and venue must appear in the body text with correct values.
    try:
        session_paragraphs = [para.text for para in doc.paragraphs
                              if 'Your session is scheduled for' in para.text]

        records_verified = 0
        for record in EXPECTED_RECORDS:
            expected_phrase = f"Your session is scheduled for {record['date']} at {record['venue']}."
            if any(expected_phrase in para_text for para_text in session_paragraphs):
                records_verified += 1

        if records_verified == 3:
            print(f"PASS: Component 3 — All 3 records have correct dates and venues (0.2 pts)")
            total_score += 0.2
        elif records_verified > 0:
            print(f"PARTIAL-FAIL: Component 3 — Only {records_verified}/3 records have correct date/venue")
        else:
            print(f"FAIL: Component 3 — No session paragraphs with correct date/venue found")
            print(f"  Found session paragraphs: {session_paragraphs}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Document has 2 page breaks separating the 3 records (0.1 points)
    # Each page break separates one merged record from the next.
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        page_break_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                for br in run._element.findall('.//w:br', ns):
                    if br.attrib.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
                    ) == 'page':
                        page_break_count += 1

        if page_break_count == 2:
            print(f"PASS: Component 4 — Exactly 2 page breaks found separating 3 records (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 4 — Expected 2 page breaks, found {page_break_count}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
