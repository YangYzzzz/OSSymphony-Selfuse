"""
Initial Setup: Mail merge invitation letter template (pre-task state)
Task ID: osworld_writer_mail_merge_003
Domain: libreoffice_writer

Creates:
  - /home/user/osworld_writer_mail_merge_003.docx  (invitation template with plain text placeholders)
  - /home/user/attendees.csv                        (data source with 10 rows)

The document has placeholder markers <<Name>>, <<EventDate>>, <<Venue>> but NO
mail merge fields connected yet. The agent's task is to connect attendees.csv
and insert the real merge fields.
"""

import csv
import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_mail_merge_003'
OUTPUT_DOCX = f'{WORKDIR}/{TASK_ID}.docx'
OUTPUT_CSV  = f'{WORKDIR}/attendees.csv'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env['DISPLAY'] = ':0'
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_csv():
    """Create the attendees data source with 10 rows."""
    rows = [
        ['Name',              'EventDate',   'Venue'],
        ['Alice Hoffman',     'April 12, 2025',  'Grand Ballroom, Ritz Hotel'],
        ['Benjamin Clarke',   'April 12, 2025',  'Grand Ballroom, Ritz Hotel'],
        ['Catherine Nguyen',  'April 13, 2025',  'Skyline Terrace, Blue Tower'],
        ['David Okafor',      'April 13, 2025',  'Skyline Terrace, Blue Tower'],
        ['Elena Vasquez',     'April 14, 2025',  'Garden Pavilion, Riverside Park'],
        ['Franklin Reyes',    'April 14, 2025',  'Garden Pavilion, Riverside Park'],
        ['Grace Kim',         'April 15, 2025',  'Conference Hall A, Metro Center'],
        ['Henry Buchanan',    'April 15, 2025',  'Conference Hall A, Metro Center'],
        ['Isabella Torres',   'April 16, 2025',  'Rooftop Lounge, The Summit'],
        ['James Whitfield',   'April 16, 2025',  'Rooftop Lounge, The Summit'],
    ]
    with open(OUTPUT_CSV, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f'CSV data source created: {OUTPUT_CSV}')


def create_docx():
    """Create the invitation letter template with plain text placeholders."""
    doc = Document()

    # --- Page margins ---
    section = doc.sections[0]
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Organisation header ---
    header_para = doc.add_paragraph()
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run('Aurora Events & Conferences')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_header.add_run('123 Lakeview Drive, Suite 400  |  events@aurora-conf.org  |  +1 800 555 0192')
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()   # spacer

    # --- Date line ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_para.add_run('March 1, 2025')

    doc.add_paragraph()   # spacer

    # --- Salutation with placeholder ---
    salutation = doc.add_paragraph()
    sal_run = salutation.add_run('Dear ')
    sal_run.font.size = Pt(11)

    placeholder_run = salutation.add_run('<<Name>>')
    placeholder_run.font.size = Pt(11)
    placeholder_run.bold = True
    placeholder_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # red to indicate placeholder

    end_run = salutation.add_run(',')
    end_run.font.size = Pt(11)

    doc.add_paragraph()   # spacer

    # --- Body paragraph 1 ---
    body1 = doc.add_paragraph()
    body1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    b1_run = body1.add_run(
        'We are delighted to extend a personal invitation to you for the '
        'Aurora Annual Innovation Summit, our flagship gathering that brings '
        'together industry leaders, visionary thinkers, and creative minds '
        'from across the globe.'
    )
    b1_run.font.size = Pt(11)

    # --- Body paragraph 2 with EventDate and Venue placeholders ---
    body2 = doc.add_paragraph()
    body2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    b2a = body2.add_run('Your session is scheduled for ')
    b2a.font.size = Pt(11)

    date_ph = body2.add_run('<<EventDate>>')
    date_ph.font.size = Pt(11)
    date_ph.bold = True
    date_ph.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    b2b = body2.add_run(' at ')
    b2b.font.size = Pt(11)

    venue_ph = body2.add_run('<<Venue>>')
    venue_ph.font.size = Pt(11)
    venue_ph.bold = True
    venue_ph.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    b2c = body2.add_run(
        '. Please ensure you have reviewed the enclosed programme and '
        'confirmed your dietary requirements by March 28, 2025.'
    )
    b2c.font.size = Pt(11)

    # --- Body paragraph 3 ---
    body3 = doc.add_paragraph()
    body3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    b3_run = body3.add_run(
        'Should you have any questions or require assistance with travel and '
        'accommodation arrangements, our dedicated concierge team is available '
        'Monday through Friday, 9 am–6 pm (GMT+8).'
    )
    b3_run.font.size = Pt(11)

    doc.add_paragraph()   # spacer

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.add_run('Warm regards,').font.size = Pt(11)

    doc.add_paragraph()   # spacer

    sig1 = doc.add_paragraph()
    s1r = sig1.add_run('Dr. Margaret Lin')
    s1r.bold = True
    s1r.font.size = Pt(11)

    sig2 = doc.add_paragraph()
    sig2.add_run('Director of Events, Aurora Events & Conferences').font.size = Pt(10)

    doc.save(OUTPUT_DOCX)
    print(f'Initial docx template created: {OUTPUT_DOCX}')


def create_initial():
    create_csv()
    create_docx()

    # GUI-ready: open the invitation template in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT_DOCX}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
