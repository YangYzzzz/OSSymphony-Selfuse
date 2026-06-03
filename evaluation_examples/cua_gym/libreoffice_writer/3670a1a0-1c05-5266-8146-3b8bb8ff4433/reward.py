"""
Reward Script: Attendance Register in LibreOffice Writer
Task ID: writer_wf_063
Domain: libreoffice_writer
Scoring:
  C1 (0.15) - Title "Attendance Register - Leadership Workshop"
  C2 (0.20) - Event details (Date, Time, Venue, Trainer)
  C3 (0.20) - Table with 5 correct header columns
  C4 (0.15) - Table has 21 rows (header + 20 numbered)
  C5 (0.10) - Total attendees line below table
  C6 (0.10) - Trainer signature line below table
  C7 (0.10) - Landscape orientation
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_063'


def persist_app_state(domain: str):
    """Try to save any unsaved GUI state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify attendance register task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph text for searching
    all_para_texts = [p.text.strip() for p in doc.paragraphs]
    all_para_texts_lower = [t.lower() for t in all_para_texts]

    # Component 1: Title "Attendance Register - Leadership Workshop" (0.15 points)
    try:
        title_found = False
        for text in all_para_texts_lower:
            if 'attendance register' in text and 'leadership workshop' in text:
                title_found = True
                break
        if title_found:
            print("PASS: Component 1 — Title contains 'Attendance Register' and 'Leadership Workshop' (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Title not found. Paragraphs: {all_para_texts[:5]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Event details - Date, Time, Venue, Trainer (0.20 points)
    # Each detail worth 0.05
    try:
        details_score = 0.0
        combined_text = ' '.join(all_para_texts_lower)

        # Also check table cells for event details
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    combined_text += ' ' + cell.text.strip().lower()

        checks = {
            'date': 'date' in combined_text,
            'time': 'time' in combined_text,
            'venue': 'venue' in combined_text,
            'trainer': 'trainer' in combined_text,
        }
        for key, found in checks.items():
            if found:
                details_score += 0.05
                print(f"  PASS: Event detail '{key}' found")
            else:
                print(f"  FAIL: Event detail '{key}' not found")

        if details_score > 0:
            print(f"PASS: Component 2 — Event details ({details_score:.2f} pts)")
            total_score += details_score
        else:
            print("FAIL: Component 2 — No event details found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table with 5 columns and correct headers (0.20 points)
    try:
        table_found = False
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_cols = len(table.columns)
            if num_cols == 5:
                # Check header row
                header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                expected_headers = ['no', 'name', 'organization', 'email', 'signature']
                matches = 0
                for exp in expected_headers:
                    for hdr in header_cells:
                        if exp in hdr:
                            matches += 1
                            break
                if matches >= 4:
                    table_found = True
                    print(f"PASS: Component 3 — Table has 5 columns with correct headers: {header_cells} ({matches}/5 match) (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"FAIL: Component 3 — Headers don't match. Found: {header_cells}, matched {matches}/5")
            else:
                print(f"FAIL: Component 3 — Table has {num_cols} columns, expected 5")
        else:
            print("FAIL: Component 3 — No table found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Table has 21 rows (header + 20 numbered rows) (0.15 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            if num_rows == 21:
                # Check that rows 1-20 are numbered
                numbered_count = 0
                for ri in range(1, min(21, num_rows)):
                    first_cell = table.rows[ri].cells[0].text.strip()
                    if first_cell.isdigit() and int(first_cell) == ri:
                        numbered_count += 1
                if numbered_count >= 18:
                    print(f"PASS: Component 4 — Table has 21 rows with {numbered_count}/20 numbered correctly (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 4 — Only {numbered_count}/20 rows numbered correctly")
            elif num_rows >= 19:
                # Partial credit: close to 21 rows
                partial = 0.15 * (num_rows / 21.0)
                partial = min(partial, 0.10)
                print(f"PARTIAL: Component 4 — Table has {num_rows} rows, expected 21 ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Table has {num_rows} rows, expected 21")
        else:
            print("FAIL: Component 4 — No table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Total attendees line (0.10 points)
    try:
        total_line_found = False
        for text in all_para_texts_lower:
            if 'total' in text and 'attend' in text:
                total_line_found = True
                break
        if total_line_found:
            print("PASS: Component 5 — Total attendees line found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 5 — No 'total attendees' line found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Trainer signature line (0.10 points)
    try:
        trainer_sig_found = False
        for text in all_para_texts_lower:
            if 'trainer' in text and 'signature' in text:
                trainer_sig_found = True
                break
        if trainer_sig_found:
            print("PASS: Component 6 — Trainer signature line found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 6 — No trainer signature line found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Landscape orientation (0.10 points)
    try:
        section = doc.sections[0]
        is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
        if not is_landscape:
            # Also check by dimensions
            is_landscape = section.page_width > section.page_height
        if is_landscape:
            print("PASS: Component 7 — Page is landscape orientation (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 — Page is portrait (width={section.page_width}, height={section.page_height})")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved state before verifying
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
