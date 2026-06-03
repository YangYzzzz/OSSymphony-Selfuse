"""
Initial Setup: Lease agreement document without watermark
Task ID: writer_legal_027
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---- Title ----
    title = doc.add_heading('RESIDENTIAL LEASE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Preamble ----
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(12)
    run = preamble.add_run(
        'This Residential Lease Agreement ("Agreement") is entered into as of '
        'March 15, 2026, by and between the following parties:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 1: Parties ----
    h1 = doc.add_heading('1. PARTIES', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('Landlord: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        'Westfield Property Management LLC, a limited liability company organized '
        'under the laws of the State of California, with its principal office at '
        '4520 Meridian Avenue, Suite 310, San Jose, CA 95124. '
        'Contact: Elena Vasquez, Property Manager, (408) 555-2847.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run = p2.add_run('Tenant: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run = p2.add_run(
        'Marcus T. Richardson, an individual residing at 1287 Oakwood Drive, '
        'Apartment 4B, Sunnyvale, CA 94086. Social Security Number ending in ***-**-4729. '
        'Contact: (408) 555-9163.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 2: Premises ----
    doc.add_heading('2. PREMISES', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'The Landlord hereby leases to the Tenant the residential property located at '
        '782 Sunridge Court, Unit 12, Mountain View, CA 94043 (the "Premises"). '
        'The Premises consists of a two-bedroom, one-bathroom apartment of approximately '
        '1,150 square feet, including a designated parking space (#34) in the underground '
        'garage and access to common amenities including the swimming pool, fitness center, '
        'and community laundry facilities.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 3: Term ----
    doc.add_heading('3. TERM OF LEASE', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'The lease term shall commence on April 1, 2026 and shall terminate on '
        'March 31, 2027 (the "Initial Term"), unless sooner terminated in accordance '
        'with the provisions of this Agreement. Upon expiration of the Initial Term, '
        'this Agreement shall automatically convert to a month-to-month tenancy under '
        'the same terms and conditions, unless either party provides at least thirty (30) '
        'days written notice of intent to terminate.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 4: Rent ----
    doc.add_heading('4. RENT', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'The Tenant agrees to pay a monthly rent of Two Thousand Eight Hundred Fifty '
        'Dollars ($2,850.00), due on the first day of each calendar month. Rent shall be '
        'paid via electronic transfer to Westfield Property Management LLC, account details '
        'to be provided separately. A late fee of $75.00 shall be assessed for any rent '
        'payment received after the fifth (5th) day of the month. If rent remains unpaid '
        'for fifteen (15) or more days, an additional late charge of $150.00 shall apply.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 5: Security Deposit ----
    doc.add_heading('5. SECURITY DEPOSIT', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Upon execution of this Agreement, the Tenant shall deposit with the Landlord '
        'the sum of Five Thousand Seven Hundred Dollars ($5,700.00) as a security deposit. '
        'The security deposit shall be held in accordance with California Civil Code '
        'Section 1950.5. The deposit, or any portion thereof, may be applied by the Landlord '
        'toward: (a) unpaid rent; (b) repair of damages beyond normal wear and tear; '
        '(c) cleaning costs to restore the Premises to the condition at move-in; or '
        '(d) other charges as permitted by law. The balance of the deposit shall be returned '
        'within twenty-one (21) days after the Tenant vacates the Premises, accompanied by '
        'an itemized statement of any deductions.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 6: Maintenance and Repairs ----
    doc.add_heading('6. MAINTENANCE AND REPAIRS', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'The Landlord shall be responsible for maintaining the structural components of the '
        'Premises, including the roof, exterior walls, plumbing, electrical systems, and '
        'heating/cooling systems in good working order. The Tenant shall be responsible for: '
        '(a) keeping the interior of the Premises clean and sanitary; (b) promptly reporting '
        'any maintenance issues or needed repairs to the Landlord; (c) not making any '
        'alterations, additions, or improvements without prior written consent of the Landlord; '
        'and (d) repairing any damage caused by the Tenant, Tenant\'s guests, or pets at '
        'the Tenant\'s expense.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 7: Utilities ----
    doc.add_heading('7. UTILITIES AND SERVICES', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'The Landlord shall be responsible for payment of the following utilities: water, '
        'sewer, and trash collection. The Tenant shall be responsible for payment of all '
        'other utilities and services, including but not limited to: electricity, natural gas, '
        'internet, cable television, and telephone. The Tenant shall establish all utility '
        'accounts in the Tenant\'s name within five (5) business days of the lease commencement '
        'date and shall maintain said accounts in good standing throughout the term of this lease.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 8: Use of Premises ----
    doc.add_heading('8. USE OF PREMISES', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'The Premises shall be used solely as a private residential dwelling for the Tenant '
        'and approved occupants. The Tenant shall not use the Premises for any commercial, '
        'illegal, or hazardous purpose. The Tenant shall comply with all applicable laws, '
        'ordinances, rules, and regulations of governmental authorities. The maximum number '
        'of occupants shall not exceed two (2) persons without prior written approval from '
        'the Landlord. Overnight guests staying more than seven (7) consecutive days or more '
        'than fourteen (14) days in any calendar month shall require Landlord\'s written consent.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 9: Pets ----
    doc.add_heading('9. PET POLICY', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Pets are permitted on the Premises subject to the following conditions: (a) a '
        'non-refundable pet deposit of $500.00 and monthly pet rent of $50.00 per pet; '
        '(b) maximum of one (1) pet, not to exceed 35 pounds at maturity; (c) the pet must '
        'be current on all vaccinations as required by local law; (d) the Tenant shall be '
        'liable for all damages caused by the pet; and (e) the Landlord reserves the right '
        'to revoke pet privileges if the pet causes disturbances or damage to the property. '
        'Prohibited breeds include Pit Bulls, Rottweilers, Doberman Pinschers, and wolf hybrids.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Section 10: Termination ----
    doc.add_heading('10. TERMINATION', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        'Either party may terminate this Agreement upon thirty (30) days written notice. '
        'The Landlord may terminate this Agreement immediately upon: (a) non-payment of rent '
        'for a period exceeding fifteen (15) days; (b) material breach of any term of this '
        'Agreement; (c) illegal activity on the Premises; or (d) actions that endanger the '
        'health and safety of other tenants. Upon termination, the Tenant shall vacate the '
        'Premises, return all keys and access devices, and leave the Premises in clean, '
        'good condition, normal wear and tear excepted.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # ---- Signature block ----
    doc.add_paragraph()  # spacer
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(24)
    run = sig.add_run('SIGNATURES')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run('_________________________________\nElena Vasquez\nWestfield Property Management LLC\nDate: _______________')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run('_________________________________\nMarcus T. Richardson\nTenant\nDate: _______________')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
