"""
Reward Script: Performance Review Form Table
Task ID: writer_hr_027
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table exists with 8 rows and 3 columns
  Component 2 (0.20): Header row contains correct column names
  Component 3 (0.30): Six competency rows have correct labels in column 0
  Component 4 (0.25): Last row is merged across all columns with 'Overall Rating and Summary'
"""

import os

from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_027'


def persist_app_state(domain):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print("PERSIST: ctrl+s sent for %s" % domain)
        except Exception as e:
            print("PERSIST_WARN: save hook failed: %s" % e)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table has correct dimensions — 8 rows, 3 columns (0.25 points)
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows == 8 and num_cols == 3:
            print("PASS: Component 1 — Table has 8 rows and 3 columns (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 1 — Expected 8 rows x 3 cols, found %d rows x %d cols" % (num_rows, num_cols))
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # Component 2: Header row contains correct column names (0.20 points)
    try:
        expected_headers = ['Competency Area', 'Rating (1-5)', 'Comments']
        actual_headers = [table.cell(0, c).text.strip() for c in range(min(3, len(table.columns)))]
        matches = sum(1 for exp, act in zip(expected_headers, actual_headers) if exp.lower() == act.lower())
        if matches == 3:
            print("PASS: Component 2 — Header row matches: %s (0.20 pts)" % actual_headers)
            total_score += 0.20
        else:
            print("FAIL: Component 2 — Expected headers %s, found %s (%d/3 match)" % (expected_headers, actual_headers, matches))
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # Component 3: Six competency rows have correct labels (0.30 points)
    # Each correct competency label is worth 0.05 points
    try:
        expected_competencies = [
            'Communication',
            'Teamwork',
            'Technical Skills',
            'Leadership',
            'Problem Solving',
            'Time Management',
        ]
        comp3_score = 0.0
        for i, expected in enumerate(expected_competencies):
            row_idx = i + 1  # rows 1-6
            if row_idx < len(table.rows):
                actual = table.cell(row_idx, 0).text.strip()
                if actual.lower() == expected.lower():
                    comp3_score += 0.05
                    print("PASS: Component 3.%d — Row %d has '%s'" % (i + 1, row_idx, actual))
                else:
                    print("FAIL: Component 3.%d — Row %d expected '%s', found '%s'" % (i + 1, row_idx, expected, actual))
            else:
                print("FAIL: Component 3.%d — Row %d does not exist" % (i + 1, row_idx))
        if comp3_score > 0:
            total_score += comp3_score
            print("Component 3 subtotal: %.2f/0.30 pts" % comp3_score)
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    # Component 4: Last row is merged across all columns with correct text (0.25 points)
    try:
        last_row_idx = len(table.rows) - 1
        if last_row_idx < 1:
            print("FAIL: Component 4 — Table has insufficient rows for summary row")
        else:
            last_row_text = table.cell(last_row_idx, 0).text.strip()

            # Check if merged: all cells in last row should reference the same text
            all_cells_same = all(
                table.cell(last_row_idx, c).text.strip() == last_row_text
                for c in range(min(3, len(table.columns)))
            )

            # Also check gridSpan in XML to confirm merge
            tc = table.cell(last_row_idx, 0)._tc
            tcPr = tc.find(qn('w:tcPr'))
            grid_span_val = 1
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    try:
                        grid_span_val = int(gs.get(qn('w:val')))
                    except (TypeError, ValueError):
                        pass

            is_merged = grid_span_val >= 3 or all_cells_same
            text_correct = 'overall rating' in last_row_text.lower() and 'summary' in last_row_text.lower()

            if is_merged and text_correct:
                print("PASS: Component 4 — Last row merged (gridSpan=%d) with text '%s' (0.25 pts)" % (grid_span_val, last_row_text))
                total_score += 0.25
            elif text_correct and not is_merged:
                print("PARTIAL: Component 4 — Text correct but not merged (0.10 pts)")
                total_score += 0.10
            elif is_merged and not text_correct:
                print("PARTIAL: Component 4 — Merged but text wrong: '%s' (0.10 pts)" % last_row_text)
                total_score += 0.10
            else:
                print("FAIL: Component 4 — Not merged and text wrong: '%s'" % last_row_text)
    except Exception as e:
        print("ERROR: Component 4 — %s" % e)

    final_score = round(min(total_score, 1.0), 2)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = '%s/%s.docx' % (WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
