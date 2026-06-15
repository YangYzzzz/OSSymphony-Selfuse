"""
Initial Setup: Easter Service Program - Unformatted initial state
Task ID: writer_creative_064
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'easter_service_program'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Remove default empty paragraph if present
    # (python-docx starts with one empty paragraph)

    def add_plain(text, size_pt=12):
        """Add a plain, left-aligned paragraph at 12pt, no formatting."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(size_pt)
        run.bold = False
        run.italic = False
        run.underline = False
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return para

    # Church name — plain, no formatting
    add_plain('Grace Community Church')

    # Date line
    add_plain('Easter Sunday Service \u2014 April 20, 2026')

    # Time line
    add_plain('10:00 AM')

    # Section heading: Order of Service — plain
    add_plain('Order of Service')

    # Service items — numbered by content (plain text, not list style)
    service_items = [
        '1. Welcome & Opening Prayer',
        '2. Hymn: Christ the Lord Is Risen Today',
        '3. Scripture Reading: John 20:1-18',
        "4. Children's Message",
        '5. Hymn: Because He Lives',
        '6. Sermon: \u201cThe Promise of Easter\u201d \u2014 Pastor David Kim',
        '7. Offering',
        '8. Hymn: Amazing Grace',
        '9. Benediction',
    ]
    for item in service_items:
        add_plain(item)

    # Section heading: Hymns — plain
    add_plain('Hymns')

    # Hymn entries — plain, no italic
    hymns = [
        'Christ the Lord Is Risen Today \u2014 #364',
        'Because He Lives \u2014 #215',
        'Amazing Grace \u2014 #185',
    ]
    for hymn in hymns:
        add_plain(hymn)

    # Section heading: Announcements — plain
    add_plain('Announcements')

    # Announcement paragraphs — realistic church announcements
    announcements = [
        (
            'Our annual Easter egg hunt will be held after the service in the church courtyard. '
            'All children ages 3\u201312 are welcome to participate. Baskets will be provided. '
            'Parents, please meet your children by the main entrance when the service concludes.'
        ),
        (
            'The women\u2019s Bible study group will resume on Tuesday, April 22, at 10:00 AM in '
            'Fellowship Hall. The new study on the Gospel of Luke begins this week. All women of '
            'the congregation are warmly invited to join.'
        ),
        (
            'Grace Community Church is collecting non-perishable food items for the local food pantry '
            'throughout the month of April. Collection bins are located in the main foyer and the '
            'south entrance. Thank you for your generous support of our neighbors in need.'
        ),
    ]
    for ann in announcements:
        add_plain(ann)

    # Remove the default empty first paragraph that python-docx adds
    # (it's at index 0 if we didn't use it)
    # We used add_paragraph() which appends, so check if first para is empty
    if doc.paragraphs and doc.paragraphs[0].text == '':
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
