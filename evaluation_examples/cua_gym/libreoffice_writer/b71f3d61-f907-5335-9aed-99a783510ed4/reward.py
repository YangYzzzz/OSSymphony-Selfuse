"""
Reward Script: Easter Service Program Formatting
Task ID: writer_creative_064
Domain: libreoffice_writer
Scoring:
  - Component 1: Church name formatted (20pt, bold, centered)           — 0.25 pts
  - Component 2: Section headers formatted (16pt, bold, underline, centered) — 0.25 pts
  - Component 3: Service items center-aligned                            — 0.20 pts
  - Component 4: Hymn entries italic and centered                        — 0.15 pts
  - Component 5: Announcement text reduced to 11pt                       — 0.15 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_064'
FILE_PATH = f'{WORKDIR}/Desktop/easter_service_program.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build a text->paragraph index for easy lookup
    para_map = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            para_map[text] = para
            # Also store by prefix for partial matches
            para_map[i] = para

    paragraphs = doc.paragraphs

    # Helper: check alignment is CENTER
    def is_centered(para):
        return para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

    # Helper: get font size in pt for first run
    def get_size_pt(para):
        for run in para.runs:
            if run.font.size is not None:
                return run.font.size.pt
        return None

    # Helper: check bold on first run
    def is_bold(para):
        for run in para.runs:
            return run.bold is True
        return False

    # Helper: check underline on first run
    def is_underlined(para):
        for run in para.runs:
            return run.underline is True
        return False

    # Helper: check italic on first run
    def is_italic(para):
        for run in para.runs:
            return run.italic is True
        return False

    # ------------------------------------------------------------------
    # Component 1: Church name formatting (20pt, bold, centered) — 0.25 pts
    # Initial: 12pt, not bold, left-aligned
    # Golden: 20pt, bold, centered
    # ------------------------------------------------------------------
    try:
        church_para = None
        for para in paragraphs:
            if 'Grace Community Church' in para.text:
                church_para = para
                break

        if church_para is None:
            print("FAIL: Component 1 — 'Grace Community Church' paragraph not found")
        else:
            size_ok = get_size_pt(church_para) == 20.0
            bold_ok = is_bold(church_para)
            centered_ok = is_centered(church_para)

            if size_ok and bold_ok and centered_ok:
                print(f"PASS: Component 1 — Church name is 20pt, bold, centered (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — Church name: size={get_size_pt(church_para)}pt (need 20), "
                      f"bold={bold_ok} (need True), centered={centered_ok} (need True)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ------------------------------------------------------------------
    # Component 2: Section headers formatted (16pt, bold, underline, centered) — 0.25 pts
    # Initial: 12pt, not bold, not underlined, left-aligned
    # Golden: 16pt, bold, underlined, centered
    # Each of the 3 section headers must pass; partial credit at 1/3 increments
    # ------------------------------------------------------------------
    try:
        section_headers = ['Order of Service', 'Hymns', 'Announcements']
        headers_passed = 0

        for header_text in section_headers:
            header_para = None
            for para in paragraphs:
                if para.text.strip() == header_text:
                    header_para = para
                    break

            if header_para is None:
                print(f"FAIL: Component 2 — '{header_text}' paragraph not found")
                continue

            size_ok = get_size_pt(header_para) == 16.0
            bold_ok = is_bold(header_para)
            underline_ok = is_underlined(header_para)
            centered_ok = is_centered(header_para)

            if size_ok and bold_ok and underline_ok and centered_ok:
                print(f"PASS: Component 2 — '{header_text}' is 16pt, bold, underlined, centered")
                headers_passed += 1
            else:
                print(f"FAIL: Component 2 — '{header_text}': size={get_size_pt(header_para)}pt (need 16), "
                      f"bold={bold_ok}, underline={underline_ok}, centered={centered_ok}")

        if headers_passed == 3:
            print(f"PASS: Component 2 — All 3 section headers correctly formatted (0.25 pts)")
            total_score += 0.25
        elif headers_passed >= 2:
            print(f"PARTIAL: Component 2 — {headers_passed}/3 headers correctly formatted (0.17 pts)")
            total_score += 0.17
        elif headers_passed == 1:
            print(f"PARTIAL: Component 2 — {headers_passed}/3 headers correctly formatted (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 2 — No section headers correctly formatted")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ------------------------------------------------------------------
    # Component 3: Service items center-aligned — 0.20 pts
    # Initial: all service items left-aligned
    # Golden: all 9 service items centered
    # ------------------------------------------------------------------
    try:
        # Service items start with "1." through "9."
        service_items = []
        for para in paragraphs:
            text = para.text.strip()
            # Match numbered items 1-9
            if len(text) >= 3 and text[0].isdigit() and text[1] == '.':
                service_items.append(para)

        if not service_items:
            print("FAIL: Component 3 — No numbered service items found")
        else:
            centered_count = sum(1 for p in service_items if is_centered(p))
            total_items = len(service_items)

            if centered_count == total_items and total_items >= 9:
                print(f"PASS: Component 3 — All {total_items} service items are centered (0.20 pts)")
                total_score += 0.20
            elif centered_count >= 7:
                print(f"PARTIAL: Component 3 — {centered_count}/{total_items} service items centered (0.14 pts)")
                total_score += 0.14
            elif centered_count >= 4:
                print(f"PARTIAL: Component 3 — {centered_count}/{total_items} service items centered (0.07 pts)")
                total_score += 0.07
            else:
                print(f"FAIL: Component 3 — Only {centered_count}/{total_items} service items centered")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ------------------------------------------------------------------
    # Component 4: Hymn entries italic and centered — 0.15 pts
    # Initial: 12pt, not italic, left-aligned
    # Golden: 12pt, italic, centered
    # Three hymn entries must be italic and centered
    # ------------------------------------------------------------------
    try:
        hymn_entries = [
            'Christ the Lord Is Risen Today — #364',
            'Because He Lives — #215',
            'Amazing Grace — #185',
        ]
        hymns_passed = 0

        for hymn_text in hymn_entries:
            hymn_para = None
            for para in paragraphs:
                if hymn_text in para.text:
                    hymn_para = para
                    break

            if hymn_para is None:
                print(f"FAIL: Component 4 — Hymn '{hymn_text[:30]}' not found")
                continue

            italic_ok = is_italic(hymn_para)
            centered_ok = is_centered(hymn_para)

            if italic_ok and centered_ok:
                print(f"PASS: Component 4 — Hymn '{hymn_text[:30]}' is italic and centered")
                hymns_passed += 1
            else:
                print(f"FAIL: Component 4 — Hymn '{hymn_text[:30]}': italic={italic_ok}, centered={centered_ok}")

        if hymns_passed == 3:
            print(f"PASS: Component 4 — All 3 hymn entries italic and centered (0.15 pts)")
            total_score += 0.15
        elif hymns_passed >= 2:
            print(f"PARTIAL: Component 4 — {hymns_passed}/3 hymns italic+centered (0.10 pts)")
            total_score += 0.10
        elif hymns_passed == 1:
            print(f"PARTIAL: Component 4 — {hymns_passed}/3 hymns italic+centered (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — No hymn entries are italic and centered")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ------------------------------------------------------------------
    # Component 5: Announcement text reduced to 11pt — 0.15 pts
    # Initial: announcement paragraphs are 12pt
    # Golden: announcement paragraphs reduced to 11pt
    # ------------------------------------------------------------------
    try:
        # Announcements section: find paragraph index for 'Announcements'
        annc_idx = None
        for i, para in enumerate(paragraphs):
            if para.text.strip() == 'Announcements':
                annc_idx = i
                break

        if annc_idx is None:
            print("FAIL: Component 5 — 'Announcements' paragraph not found")
        else:
            # Announcement text paragraphs come after 'Announcements'
            annc_paras = []
            for para in paragraphs[annc_idx + 1:]:
                text = para.text.strip()
                if text:
                    annc_paras.append(para)

            if not annc_paras:
                print("FAIL: Component 5 — No announcement text paragraphs found after 'Announcements'")
            else:
                size_11_count = sum(1 for p in annc_paras if get_size_pt(p) == 11.0)
                total_annc = len(annc_paras)

                if size_11_count == total_annc and total_annc >= 3:
                    print(f"PASS: Component 5 — All {total_annc} announcement paragraphs are 11pt (0.15 pts)")
                    total_score += 0.15
                elif size_11_count >= 2:
                    print(f"PARTIAL: Component 5 — {size_11_count}/{total_annc} announcement paragraphs are 11pt (0.10 pts)")
                    total_score += 0.10
                elif size_11_count == 1:
                    print(f"PARTIAL: Component 5 — {size_11_count}/{total_annc} announcement paragraphs are 11pt (0.05 pts)")
                    total_score += 0.05
                else:
                    print(f"FAIL: Component 5 — Announcement paragraphs: sizes={[get_size_pt(p) for p in annc_paras]} (expected 11pt)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
