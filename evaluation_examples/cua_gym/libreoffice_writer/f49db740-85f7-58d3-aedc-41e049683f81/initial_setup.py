"""
Initial Setup: Create a Writer document with contract clauses for auto-text entry creation task
Task ID: writer_biz_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_078'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

INDEMNITY_CLAUSE = (
    "Indemnification Clause. The Contractor shall indemnify, defend, and hold harmless the Company, "
    "its officers, directors, employees, agents, and affiliates from and against any and all claims, "
    "damages, losses, liabilities, costs, and expenses (including reasonable attorneys' fees) arising "
    "out of or related to: (a) any breach of this Agreement by the Contractor; (b) any negligent or "
    "wrongful act or omission of the Contractor or its personnel in the performance of services under "
    "this Agreement; (c) any violation of applicable law by the Contractor; or (d) any infringement "
    "of third-party intellectual property rights resulting from the Contractor's deliverables. This "
    "indemnification obligation shall survive the termination or expiration of this Agreement for a "
    "period of three (3) years."
)

CONFIDENTIALITY_CLAUSE = (
    "Confidentiality Clause. Each party acknowledges that in connection with the performance of this "
    "Agreement, it may receive or have access to Confidential Information of the other party. "
    "'Confidential Information' means all non-public information disclosed by either party, whether "
    "orally, in writing, or by inspection, including but not limited to trade secrets, business plans, "
    "financial data, customer lists, technical specifications, and proprietary software. The receiving "
    "party agrees to: (i) maintain all Confidential Information in strict confidence; (ii) not disclose "
    "Confidential Information to any third party without prior written consent; (iii) use Confidential "
    "Information solely for the purposes of this Agreement; and (iv) return or destroy all Confidential "
    "Information upon termination of this Agreement. These obligations shall remain in effect for five "
    "(5) years from the date of disclosure."
)


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    heading = doc.add_heading("Standard Contract Clauses", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Indemnification clause paragraph
    para1 = doc.add_paragraph()
    run1 = para1.add_run(INDEMNITY_CLAUSE)
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(12)
    para1.paragraph_format.space_after = Pt(12)

    # Confidentiality clause paragraph
    para2 = doc.add_paragraph()
    run2 = para2.add_run(CONFIDENTIALITY_CLAUSE)
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(12)
    para2.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
