"""
Initial Setup: Language exercise document with 5 paragraphs of body text, all in black.
Task ID: osworld_writer_vowel_consonant_coloring_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("Language Exercise: Vowel and Consonant Identification", level=0)

    # Subtitle / instruction paragraph
    subtitle = doc.add_paragraph(
        "Instructions: Read the following passages carefully. "
        "Identify each word by its starting letter: vowel or consonant."
    )
    subtitle.paragraph_format.space_after = Pt(12)

    # 5 body paragraphs — all text in plain black, no color formatting
    paragraphs_text = [
        # Paragraph 1
        (
            "Every morning, Anna observed remarkable transformations outside her window. "
            "Ancient oak trees swayed gently in soft breezes, their emerald leaves catching "
            "golden sunlight. Numerous birds arrived early, singing melodious tunes across "
            "open fields. Animals wandered peacefully, exploring endless meadows beneath "
            "a brilliant azure sky. Often, she appreciated these unforgettable natural moments."
        ),
        # Paragraph 2
        (
            "Benjamin studied diligently at a prestigious university, pursuing advanced degrees "
            "in organic chemistry. Each experiment required unusual precision, involving "
            "intricate measurements and careful observation. His professor often emphasized "
            "analytical thinking over rote memorization. Understanding complex reactions gave "
            "Benjamin enormous satisfaction, inspiring ambitious goals for innovative research."
        ),
        # Paragraph 3
        (
            "Across urban districts, extraordinary changes transformed ordinary neighborhoods. "
            "Developers erected impressive apartment buildings along old industrial avenues. "
            "Elegant restaurants opened alongside artisan bakeries, attracting enthusiastic "
            "visitors from around entire regions. Infrastructure improvements enabled efficient "
            "transportation options, allowing residents easy access to outstanding cultural events."
        ),
        # Paragraph 4
        (
            "Outside a remote village, experienced farmers harvested abundant crops under "
            "overcast skies. Organized teams operated large equipment across extensive farmland. "
            "Ancient irrigation systems efficiently delivered essential water to arid zones. "
            "Every harvest season offered opportunities for entire communities, ensuring adequate "
            "food supplies. Automation introduced unique advantages, improving overall agricultural "
            "output enormously."
        ),
        # Paragraph 5
        (
            "International athletes assembled for an extraordinary Olympic competition, "
            "representing over eighty nations across all inhabited continents. Each athlete "
            "endured intense training, overcoming adversity and exhaustion. Organizers arranged "
            "impressive opening ceremonies, attracting enormous audiences globally. "
            "Unexpected victories inspired universal admiration, encouraging younger generations "
            "to aspire toward athletic excellence and uphold Olympic ideals."
        ),
    ]

    for para_text in paragraphs_text:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(10)
        para.paragraph_format.line_spacing = 1.15
        # Add entire paragraph as one run with black (default) color — no color coding applied
        run = para.add_run(para_text)
        run.font.size = Pt(12)
        # Explicitly set color to black (auto/default black)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
