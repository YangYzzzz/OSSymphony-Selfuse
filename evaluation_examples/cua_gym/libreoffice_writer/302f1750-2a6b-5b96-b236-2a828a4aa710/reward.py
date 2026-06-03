"""
Reward Script: Vowel/Consonant Color Coding in Language Exercise Document
Task ID: osworld_writer_vowel_consonant_coloring_008
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): At least one body paragraph has color applied (any non-black colors present)
  Component 2 (0.4): All vowel-starting words are colored red (FF0000) across all 5 body paragraphs
  Component 3 (0.3): All consonant-starting words are colored blue (0000FF) across all 5 body paragraphs
Total: 1.0
"""

import os

from docx import Document
from docx.shared import RGBColor

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_008'

VOWELS = set('aeiouAEIOU')
# Body paragraphs are indices 2-6 (0=Title, 1=Instructions, 2-6=5 body paragraphs)
BODY_PARA_INDICES = [2, 3, 4, 5, 6]

RED_TARGET = RGBColor(0xFF, 0x00, 0x00)
BLUE_TARGET = RGBColor(0x00, 0x00, 0xFF)


def color_distance(rgb1, rgb2):
    """Euclidean distance between two RGBColor values."""
    if rgb1 is None or rgb2 is None:
        return 999
    return ((rgb1[0] - rgb2[0]) ** 2 +
            (rgb1[1] - rgb2[1]) ** 2 +
            (rgb1[2] - rgb2[2]) ** 2) ** 0.5


def get_word_color(run):
    """Return the RGBColor of a run, or None if not set."""
    try:
        if run.font.color and run.font.color.rgb:
            return run.font.color.rgb
    except Exception:
        pass
    return None


def verify_task(file_path):
    """
    Verify vowel/consonant color coding task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: must have at least 7 paragraphs (Title + Instructions + 5 body)
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs; expected at least 7.")
        print("REWARD: 0.0")
        return 0.0

    # Collect all body paragraph word runs for scoring
    # Each "word run" is a run whose stripped text starts with an alphabetic character
    vowel_words = []   # (run_text, actual_rgb)
    consonant_words = []  # (run_text, actual_rgb)
    any_colored = False

    for idx in BODY_PARA_INDICES:
        para = doc.paragraphs[idx]
        for run in para.runs:
            text = run.text.strip().strip('.,;:!?()[]{}"\'-')
            if not text:
                continue
            first = text[0]
            rgb = get_word_color(run)
            if rgb is not None and str(rgb) != '000000':
                any_colored = True
            if first in VOWELS:
                vowel_words.append((run.text.strip(), rgb))
            elif first.isalpha():
                consonant_words.append((run.text.strip(), rgb))

    # Component 1: At least one body paragraph has color applied (0.3 points)
    # This distinguishes initial (all black) from any attempt at coloring
    try:
        if any_colored:
            print(f"PASS: Component 1 — Color coding applied to body paragraphs (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No color applied to any body paragraph text")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All vowel-starting words are red FF0000 (0.4 points)
    # Red = RGBColor(255, 0, 0) for vowel-starting words
    try:
        if not vowel_words:
            print("FAIL: Component 2 — No vowel-starting words found in body paragraphs")
        else:
            correct_vowels = 0
            incorrect_vowels = []
            for word, rgb in vowel_words:
                dist = color_distance(rgb, RED_TARGET)
                if dist < 50:
                    correct_vowels += 1
                else:
                    actual_str = str(rgb) if rgb else 'None'
                    incorrect_vowels.append(f"{word!r}: expected FF0000, got {actual_str}")

            pct = correct_vowels / len(vowel_words) if vowel_words else 0
            if pct == 1.0:
                print(f"PASS: Component 2 — All {len(vowel_words)} vowel-starting words are red (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — Only {correct_vowels}/{len(vowel_words)} vowel-starting words are red")
                if incorrect_vowels:
                    for item in incorrect_vowels[:5]:
                        print(f"  {item}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All consonant-starting words are blue 0000FF (0.3 points)
    # Blue = RGBColor(0, 0, 255) for consonant-starting words
    try:
        if not consonant_words:
            print("FAIL: Component 3 — No consonant-starting words found in body paragraphs")
        else:
            correct_consonants = 0
            incorrect_consonants = []
            for word, rgb in consonant_words:
                dist = color_distance(rgb, BLUE_TARGET)
                if dist < 50:
                    correct_consonants += 1
                else:
                    actual_str = str(rgb) if rgb else 'None'
                    incorrect_consonants.append(f"{word!r}: expected 0000FF, got {actual_str}")

            pct = correct_consonants / len(consonant_words) if consonant_words else 0
            if pct == 1.0:
                print(f"PASS: Component 3 — All {len(consonant_words)} consonant-starting words are blue (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 3 — Only {correct_consonants}/{len(consonant_words)} consonant-starting words are blue")
                if incorrect_consonants:
                    for item in incorrect_consonants[:5]:
                        print(f"  {item}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
