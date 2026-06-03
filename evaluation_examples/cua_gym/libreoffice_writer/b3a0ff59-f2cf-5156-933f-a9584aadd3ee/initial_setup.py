"""
Initial Setup: Create a Writer document with a two-level numbered list containing 8 items.
Task ID: writer_lec_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    doc.add_heading("Strategic Marketing Plan 2026", level=1)
    doc.add_paragraph("")  # spacer

    # Two-level numbered list: 5 level-1 items, 3 level-2 items (8 total)
    # We use "List Number" for level 1 and "List Number 2" for level 2.

    # Item 1 (level 1)
    doc.add_paragraph("Conduct comprehensive market research across all target demographics", style="List Number")

    # Item 2 (level 1)
    doc.add_paragraph("Develop brand positioning strategy for the Asia-Pacific region", style="List Number")

    # Item 2a (level 2)
    doc.add_paragraph("Identify key competitors and analyze their market share", style="List Number 2")

    # Item 2b (level 2)
    doc.add_paragraph("Define unique value propositions for each sub-market", style="List Number 2")

    # Item 3 (level 1)
    doc.add_paragraph("Design integrated digital advertising campaigns for Q3 and Q4", style="List Number")

    # Item 3a (level 2)
    doc.add_paragraph("Allocate budget across social media, search, and display channels", style="List Number 2")

    # Item 4 (level 1)
    doc.add_paragraph("Establish partnerships with regional influencers and content creators", style="List Number")

    # Item 5 (level 1)
    doc.add_paragraph("Set up quarterly performance review meetings with all stakeholders", style="List Number")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
