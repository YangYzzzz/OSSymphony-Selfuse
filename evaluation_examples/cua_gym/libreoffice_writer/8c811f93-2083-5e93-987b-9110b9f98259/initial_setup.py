"""
Initial Setup: Custom page size configuration for a marketing brochure
Task ID: writer_page_005
Domain: libreoffice_writer

Creates custom_brochure.docx on the Desktop with A4 page settings.
The user's task is to change the page size to 20cm x 25cm.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'custom_brochure'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page to A4: 21.0cm x 29.7cm, portrait
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Margins: all 2.54cm
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Page 1: Cover Page ---
    title = doc.add_heading('Horizon Living', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph('Discover Your Perfect Home')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph('')  # Spacer

    intro = doc.add_paragraph(
        'Welcome to Horizon Living — where modern design meets comfortable living. '
        'Our properties are crafted for those who appreciate quality, style, and community.'
    )
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph('')  # Spacer

    # Features list
    doc.add_heading('Featured Properties', level=2)

    properties = [
        ('Maple Grove Residences', 'A tranquil suburban community with 3–5 bedroom homes. '
         'Features landscaped gardens, community pool, and easy freeway access. '
         'Starting from $485,000.'),
        ('Skyline Tower Apartments', 'Luxury high-rise living in the heart of downtown. '
         '1–3 bedroom units with panoramic city views, concierge service, and rooftop terrace. '
         'Starting from $620,000.'),
        ('Riverside Villas', 'Exclusive waterfront villas with private docks. '
         '4–6 bedroom homes with open-plan living, chef kitchens, and smart home features. '
         'Starting from $1,250,000.'),
    ]

    for prop_name, prop_desc in properties:
        p = doc.add_paragraph(style='List Bullet')
        run_name = p.add_run(prop_name)
        run_name.bold = True
        run_name.font.size = Pt(12)
        run_desc = p.add_run(f': {prop_desc}')
        run_desc.font.size = Pt(11)

    # Page break to go to page 2
    doc.add_page_break()

    # --- Page 2: Details & Contact ---
    doc.add_heading('Why Choose Horizon Living?', level=1)

    benefits = [
        'Award-winning architecture and interior design',
        'Energy-efficient construction with solar panels',
        'Close proximity to top-rated schools and hospitals',
        'Flexible financing options with partnered banks',
        'Dedicated after-sales support and maintenance services',
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_paragraph('')  # Spacer

    doc.add_heading('Customer Testimonials', level=2)

    testimonials = [
        ('"Moving into Maple Grove was the best decision our family made. '
         'The neighborhood is safe, friendly, and beautifully maintained."',
         '— Jessica Hartwell, Maple Grove Resident'),
        ('"The team at Horizon Living made the entire purchasing process seamless. '
         'Our Skyline Tower apartment exceeded all expectations."',
         '— David Nguyen, Skyline Tower Owner'),
    ]

    for quote, attribution in testimonials:
        p_quote = doc.add_paragraph(quote)
        p_quote.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p_quote.runs[0].italic = True

        p_attr = doc.add_paragraph(attribution)
        p_attr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p_attr.runs[0].bold = True
        p_attr.runs[0].font.size = Pt(10)

        doc.add_paragraph('')  # Spacer

    doc.add_heading('Contact Us', level=2)

    contact_info = [
        'Address: 48 Meridian Boulevard, Suite 1200, Metro City, MC 90210',
        'Phone: +1 (800) 555-0192',
        'Email: sales@horizonliving.com',
        'Website: www.horizonliving.com',
        'Office Hours: Monday – Friday, 9:00 AM – 6:00 PM',
    ]
    for info in contact_info:
        doc.add_paragraph(info, style='List Bullet')

    doc.add_paragraph('')  # Spacer

    disclaimer = doc.add_paragraph(
        'All prices and availability are subject to change without notice. '
        'Images are for illustrative purposes only. '
        'Please contact our sales team for the most up-to-date information.'
    )
    disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
