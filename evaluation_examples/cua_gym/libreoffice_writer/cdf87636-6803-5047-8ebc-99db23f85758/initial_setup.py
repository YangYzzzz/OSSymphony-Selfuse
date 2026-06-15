"""
Initial Setup: Mathematics methodology document awaiting quadratic formula insertion
Task ID: writer_acad_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_026'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Mathematical Methods in Applied Sciences', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author line ---
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Dr. Elena Vasquez — Department of Applied Mathematics')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This document presents a comprehensive overview of fundamental mathematical '
        'methods used in applied sciences. We examine analytical techniques for solving '
        'polynomial equations, differential equations, and optimization problems. Special '
        'attention is given to closed-form solutions that have broad applicability across '
        'engineering, physics, and computational disciplines.'
    )

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The ability to solve polynomial equations is central to many branches of science '
        'and engineering. From determining eigenvalues in structural analysis to finding '
        'equilibrium points in chemical kinetics, algebraic solutions form the backbone '
        'of quantitative modeling. This paper reviews both classical and modern approaches '
        'to polynomial root-finding, beginning with the well-known quadratic case.'
    )
    doc.add_paragraph(
        'Throughout history, mathematicians have sought general formulas for the roots of '
        'polynomial equations. While Galois theory demonstrates that no such formula exists '
        'for polynomials of degree five or higher, exact solutions are available for '
        'quadratic, cubic, and quartic equations. Among these, the quadratic formula remains '
        'the most widely applied in practical computation.'
    )

    # --- Section 2: Methodology ---
    doc.add_heading('2. Methodology for Solving Second-Degree Equations', level=1)
    doc.add_paragraph(
        'Consider the general second-degree polynomial equation of the form '
        'ax\u00B2 + bx + c = 0, where a, b, and c are real-valued coefficients with '
        'a \u2260 0. The solution to this equation is given by the quadratic formula, '
        'which provides an explicit expression for the two roots in terms of the '
        'coefficients. The formula is presented below:'
    )

    # --- Placeholder paragraph indicating the formula should follow ---
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = placeholder.add_run('[Insert quadratic formula here using the equation editor]')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(11)

    # --- Continuation after formula ---
    doc.add_paragraph(
        'The expression under the square root sign, b\u00B2 \u2212 4ac, is known as the '
        'discriminant (\u0394). The discriminant determines the nature of the roots:'
    )

    # Bullet points about discriminant
    doc.add_paragraph(
        'If \u0394 > 0, the equation has two distinct real roots.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'If \u0394 = 0, the equation has exactly one repeated real root (a double root).',
        style='List Bullet'
    )
    doc.add_paragraph(
        'If \u0394 < 0, the equation has two complex conjugate roots.',
        style='List Bullet'
    )

    # --- Section 3: Applications ---
    doc.add_heading('3. Applications', level=1)
    doc.add_paragraph(
        'The quadratic formula finds extensive use in projectile motion calculations, '
        'circuit analysis, optimization of parabolic reflectors, and financial modeling. '
        'In projectile motion, for instance, setting the height equation h(t) = '
        '-\u00BDgt\u00B2 + v\u2080t + h\u2080 equal to zero and solving for t yields '
        'the time of flight, directly leveraging the quadratic formula.'
    )
    doc.add_paragraph(
        'In electrical engineering, the resonant frequency of an RLC circuit is determined '
        'by solving a quadratic characteristic equation. The roots of this equation reveal '
        'whether the system is overdamped, critically damped, or underdamped, each regime '
        'corresponding to qualitatively different transient behaviors.'
    )

    # --- Section 4: Numerical Examples ---
    doc.add_heading('4. Numerical Examples', level=1)

    # Table of examples
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'

    headers = ['Equation', 'a', 'b', 'c', 'Discriminant (\u0394)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    examples = [
        ['2x\u00B2 + 5x \u2212 3 = 0',   '2',  '5',  '\u22123', '49'],
        ['x\u00B2 \u2212 4x + 4 = 0',     '1',  '\u22124', '4',  '0'],
        ['3x\u00B2 + 2x + 1 = 0',          '3',  '2',  '1',      '\u22128'],
        ['x\u00B2 \u2212 7x + 10 = 0',     '1',  '\u22127', '10', '9'],
    ]
    for r, row_data in enumerate(examples, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacing

    # --- References ---
    doc.add_heading('References', level=1)
    doc.add_paragraph(
        '[1] Stewart, J. (2020). Calculus: Early Transcendentals, 9th Edition. Cengage Learning.'
    )
    doc.add_paragraph(
        '[2] Artin, M. (2018). Algebra, 2nd Edition. Pearson.'
    )
    doc.add_paragraph(
        '[3] Strang, G. (2016). Introduction to Linear Algebra, 5th Edition. '
        'Wellesley-Cambridge Press.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
