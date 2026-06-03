"""
Initial Setup: Blog post with em dashes for replacement task
Task ID: writer_edit_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'blog_post'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ---- Page 1 ----

    # Title
    title = doc.add_heading("The Digital Productivity Revolution", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / author line
    subtitle = doc.add_paragraph("By Jordan Maxwell  |  Tech & Society Blog")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph("")  # spacer

    # Introduction paragraph
    intro = doc.add_paragraph(
        "In the last decade, the way we work has shifted dramatically. "
        "Digital tools have redefined everything from communication to project management. "
        "Yet with this transformation comes a new challenge: information overload. "
        "Professionals at every level are struggling to find focus in an era of constant connectivity."
    )

    doc.add_paragraph("")  # spacer

    # Section heading
    doc.add_heading("A Brief History of Productivity Methods", level=2)

    # Paragraph with page range em dash (number range)
    doc.add_paragraph(
        "The modern productivity movement traces its roots to Frederick Taylor's work in the early 1900s. "
        "His landmark study, covered in pages 10\u201420 of the seminal text 'Scientific Management', "
        "introduced the idea of optimizing workflows for maximum efficiency. "
        "While controversial in its time, Taylor's methods laid the groundwork for today's agile frameworks."
    )

    # Paragraph with year range em dash (number range)
    doc.add_paragraph(
        "The digital era truly took hold during the years 2020\u20142025, a period marked by remote work adoption "
        "and rapid advances in cloud collaboration tools. "
        "Companies that embraced distributed teams during this window gained a significant competitive edge, "
        "forcing traditional office-centric organizations to adapt or fall behind."
    )

    doc.add_paragraph("")  # spacer

    # Section heading
    doc.add_heading("Core Principles for the Modern Worker", level=2)

    # Paragraph with parenthetical em dashes (first occurrence)
    doc.add_paragraph(
        "Among the many frameworks proposed over the years, the GTD (Getting Things Done) methodology "
        "has proven most durable. The solution \u2014 though controversial among minimalists \u2014 was effective "
        "precisely because it acknowledged the reality of modern work: we are managing not just tasks, "
        "but commitments, contexts, and energy levels simultaneously."
    )

    # Add page break to create page 2
    doc.add_page_break()

    # ---- Page 2 ----

    doc.add_heading("Deep Work and the Attention Economy", level=2)

    # Paragraph with chapter range em dash (number range)
    doc.add_paragraph(
        "Cal Newport's influential research, detailed in chapters 3\u20147 of his book 'Deep Work', "
        "argues that the ability to focus without distraction is becoming increasingly rare and increasingly valuable. "
        "Newport distinguishes between shallow work (easily replicated, low cognitive demand) "
        "and deep work, which requires sustained concentration and produces the most meaningful output."
    )

    doc.add_paragraph("")  # spacer

    doc.add_heading("Practical Strategies That Work", level=2)

    # Paragraph with second parenthetical em dashes
    doc.add_paragraph(
        "Time-blocking, a method endorsed by numerous productivity experts, involves scheduling every hour of the workday. "
        "He arrived \u2014 finally \u2014 at noon after a morning of back-to-back meetings that had left him mentally drained. "
        "That moment of clarity, sitting down to a fully blocked calendar, transformed his output in the weeks that followed."
    )

    # Concluding paragraph
    doc.add_paragraph(
        "The evidence is clear: structured, intentional work habits consistently outperform reactive, "
        "always-on approaches. Whether you adopt a formal system or build your own hybrid method, "
        "the key is consistency. Start small, track your progress, and iterate based on what works for you."
    )

    doc.add_paragraph("")  # spacer

    # Closing
    doc.add_heading("Conclusion", level=2)

    doc.add_paragraph(
        "The digital productivity revolution is not slowing down. "
        "New tools, new research, and new demands will continue to reshape how we work. "
        "But the fundamentals remain constant: clarity of purpose, disciplined execution, and regular reflection "
        "are the cornerstones of any sustainable productivity system."
    )

    # Author bio footer note
    doc.add_paragraph("")
    bio = doc.add_paragraph(
        "Jordan Maxwell is a workplace productivity consultant and regular contributor to Tech & Society Blog. "
        "Follow the series for weekly insights on digital work culture."
    )
    bio.runs[0].italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
