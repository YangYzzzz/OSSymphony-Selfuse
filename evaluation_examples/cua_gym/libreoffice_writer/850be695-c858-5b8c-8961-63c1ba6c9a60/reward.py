"""
Reward Script: Replace em dashes with en dashes (number ranges) and commas (parenthetical)
Task ID: writer_edit_041
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): All 3 number-range em dashes converted to en dashes in paras 6, 7, 13
  Component 2 (0.3): Both parenthetical em dash pairs converted to commas (paras 10, 16)
  Component 3 (0.3): No em dashes remain anywhere in the document
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_041'

EM_DASH = '\u2014'  # — em dash (original)
EN_DASH = '\u2013'  # – en dash (expected for number ranges)


def persist_app_state():
    """Try to save any open LibreOffice Writer document before scoring."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires:
    1. Em dashes between numbers (number ranges) → replaced with en dashes
       - 'pages 10—20'   → 'pages 10–20'     (para 6)
       - 'years 2020—2025' → 'years 2020–2025' (para 7)
       - 'chapters 3—7'  → 'chapters 3–7'    (para 13)
    2. Em dashes used as parenthetical separators → replaced with commas (spaces removed)
       - 'The solution — though controversial ... — was effective'
         → 'The solution, though controversial, was effective'  (para 10)
       - 'He arrived — finally — at noon'
         → 'He arrived, finally, at noon'                      (para 16)
    3. No em dashes remain in the document at all.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build full-text per paragraph for easy access
    para_texts = [p.text for p in doc.paragraphs]

    # -----------------------------------------------------------------------
    # Component 1: Number-range em dashes converted to en dashes (0.4 points)
    # All three number ranges must use en dashes, not em dashes.
    # Para 6: 'pages 10–20'
    # Para 7: 'years 2020–2025'
    # Para 13: 'chapters 3–7'
    # -----------------------------------------------------------------------
    try:
        checks = [
            ("pages 10" + EN_DASH + "20",   "pages 10\u201420",  "pages 10–20",    6),
            ("years 2020" + EN_DASH + "2025", "years 2020\u20142025", "years 2020–2025", 7),
            ("chapters 3" + EN_DASH + "7",  "chapters 3\u20147",  "chapters 3–7",  13),
        ]
        passed = 0
        for expected_en, expected_em, label, para_idx in checks:
            text = para_texts[para_idx] if para_idx < len(para_texts) else ""
            if expected_en in text:
                print(f"PASS: Component 1 — '{label}' uses en dash correctly in para {para_idx}")
                passed += 1
            elif expected_em in text:
                print(f"FAIL: Component 1 — '{label}' still uses em dash in para {para_idx}")
            else:
                print(f"FAIL: Component 1 — '{label}' not found at expected location (para {para_idx}); text snippet: {repr(text[:80])}")

        if passed == 3:
            print(f"PASS: Component 1 — All 3 number-range en dashes present (0.4 pts)")
            total_score += 0.4
        elif passed > 0:
            partial = round(passed / 3 * 0.4, 4)
            print(f"PARTIAL: Component 1 — {passed}/3 number-range en dashes present ({partial} pts)")
            total_score += partial
        else:
            print("FAIL: Component 1 — No number-range en dashes found (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Parenthetical em dashes replaced with commas (0.3 points)
    # Para 10: 'The solution, though controversial, was effective'
    # Para 16: 'He arrived, finally, at noon'
    # Also verify that no em dash remains in these paragraphs.
    # -----------------------------------------------------------------------
    try:
        p10_text = para_texts[10] if len(para_texts) > 10 else ""
        p16_text = para_texts[16] if len(para_texts) > 16 else ""

        # Expected golden substrings
        p10_expected = "The solution, though controversial, was effective"
        p16_expected = "He arrived, finally, at noon"

        # Em dash versions (must NOT appear in golden)
        p10_em = "The solution " + EM_DASH
        p16_em = "He arrived " + EM_DASH

        p10_pass = (p10_expected in p10_text) and (EM_DASH not in p10_text)
        p16_pass = (p16_expected in p16_text) and (EM_DASH not in p16_text)

        if p10_pass:
            print(f"PASS: Component 2 — Para 10 parenthetical correctly uses commas")
        elif p10_em in p10_text:
            print(f"FAIL: Component 2 — Para 10 still contains em dash (not replaced with commas)")
        else:
            print(f"FAIL: Component 2 — Para 10 expected pattern not found; snippet: {repr(p10_text[:100])}")

        if p16_pass:
            print(f"PASS: Component 2 — Para 16 parenthetical correctly uses commas")
        elif p16_em in p16_text:
            print(f"FAIL: Component 2 — Para 16 still contains em dash (not replaced with commas)")
        else:
            print(f"FAIL: Component 2 — Para 16 expected pattern not found; snippet: {repr(p16_text[:100])}")

        if p10_pass and p16_pass:
            print("PASS: Component 2 — Both parenthetical replacements correct (0.3 pts)")
            total_score += 0.3
        elif p10_pass or p16_pass:
            print("PARTIAL: Component 2 — One of two parenthetical replacements correct (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 2 — Neither parenthetical replacement correct (0.0 pts)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: No em dashes remain anywhere in the document (0.3 points)
    # After the task, all 7 original em dashes should be replaced.
    # -----------------------------------------------------------------------
    try:
        remaining_em = []
        for i, para in enumerate(doc.paragraphs):
            if EM_DASH in para.text:
                remaining_em.append((i, para.text[:100]))

        if not remaining_em:
            print("PASS: Component 3 — No em dashes remain in document (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — {len(remaining_em)} paragraph(s) still contain em dashes:")
            for idx, snippet in remaining_em:
                print(f"  Para {idx}: {repr(snippet)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 4), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path in given env
file_path = f'{WORKDIR}/Desktop/blog_post.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
