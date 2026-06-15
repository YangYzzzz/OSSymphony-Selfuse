"""
Initial Setup: Set up different first page formatting
Task ID: writer_biz_035
Domain: libreoffice_writer

Creates an 8-page business document with Default Page Style applied uniformly.
All pages (including the first) have:
  - Header: "Meridian Solutions Inc."
  - Footer: page numbers
No "First Page" style differentiation exists yet.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_035'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_to_footer(section):
    """Add a PAGE field code to the footer of the given section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_prefix = fp.add_run("Page ")
    run_prefix.font.size = Pt(10)
    run_prefix.font.name = "Liberation Sans"

    # PAGE field: begin
    r1 = fp.add_run()
    r1.font.size = Pt(10)
    r1.font.name = "Liberation Sans"
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    # PAGE field: instruction
    r2 = fp.add_run()
    r2.font.size = Pt(10)
    r2.font.name = "Liberation Sans"
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    instr.set(qn('xml:space'), 'preserve')
    r2._element.append(instr)

    # PAGE field: end
    r3 = fp.add_run()
    r3.font.size = Pt(10)
    r3.font.name = "Liberation Sans"
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def set_header(section, text):
    """Set header text for the given section."""
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = hp.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Liberation Sans"
    run.bold = True


def create_initial():
    doc = Document()

    # --- Page/Section Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Ensure "different first page" is OFF so all pages share the same header/footer
    sectPr = section._sectPr
    # Remove titlePg element if present (this disables "different first page")
    for tp in sectPr.findall(qn('w:titlePg')):
        sectPr.remove(tp)

    # Set header and footer on the default section (applies to ALL pages)
    set_header(section, "Meridian Solutions Inc.")
    add_page_number_to_footer(section)

    # --- Page 1: Company Overview ---
    h = doc.add_heading("Meridian Solutions Inc. — Annual Business Review", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Meridian Solutions Inc. is a mid-sized technology consulting firm headquartered "
        "in Austin, Texas. Founded in 2011, the company has grown to serve over 200 enterprise "
        "clients across North America, Europe, and the Asia-Pacific region. Our core competencies "
        "include cloud migration, enterprise architecture, cybersecurity advisory, and digital "
        "transformation strategy."
    )
    doc.add_paragraph(
        "This document provides a comprehensive overview of the company's financial performance, "
        "departmental achievements, strategic initiatives, and outlook for the upcoming fiscal year. "
        "All data referenced herein pertains to the fiscal year ending December 31, 2025."
    )
    doc.add_paragraph(
        "Our leadership team, consisting of 12 senior executives, continues to drive innovation "
        "through a balanced approach of organic growth and strategic acquisitions. In Q3 2025, "
        "Meridian acquired DataBridge Analytics, a data engineering firm based in Toronto, expanding "
        "our capabilities in real-time data processing and machine learning operations."
    )

    # --- Page 2: Financial Performance ---
    doc.add_page_break()
    doc.add_heading("Financial Performance — FY2025", level=1)
    doc.add_paragraph(
        "Total revenue for FY2025 reached $187.4 million, representing a 14.2% year-over-year "
        "increase from $164.1 million in FY2024. Operating income improved to $31.8 million, "
        "driven by higher-margin consulting engagements and the successful integration of recently "
        "acquired business units."
    )

    # Revenue table
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Quarter", "Revenue ($M)", "Operating Income ($M)", "Margin (%)"]
    for i, h_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ["Q1 2025", "$41.2", "$6.8", "16.5%"],
        ["Q2 2025", "$44.7", "$7.4", "16.6%"],
        ["Q3 2025", "$49.1", "$8.6", "17.5%"],
        ["Q4 2025", "$52.4", "$9.0", "17.2%"],
        ["FY2025 Total", "$187.4", "$31.8", "17.0%"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph("")
    doc.add_paragraph(
        "Key financial highlights include a 22% increase in recurring revenue from managed services "
        "contracts and a significant reduction in client acquisition costs through improved digital "
        "marketing campaigns. The EBITDA margin expanded by 1.8 percentage points compared to FY2024."
    )

    # --- Page 3: Department Highlights ---
    doc.add_page_break()
    doc.add_heading("Department Highlights", level=1)

    doc.add_heading("Cloud Infrastructure Division", level=2)
    doc.add_paragraph(
        "Led by VP Sarah Chen, the Cloud Infrastructure Division delivered 68 enterprise migration "
        "projects in FY2025, up from 51 in the prior year. Notable engagements included the full "
        "AWS migration for Continental Banking Group and the hybrid cloud deployment for Nexus "
        "Pharmaceuticals. Division revenue: $62.3 million."
    )

    doc.add_heading("Cybersecurity Advisory", level=2)
    doc.add_paragraph(
        "Under Director Marcus Johnson, the Cybersecurity Advisory practice expanded its penetration "
        "testing services and launched a new Managed Detection and Response (MDR) offering. The team "
        "conducted 124 security assessments and responded to 18 incident response engagements. "
        "Division revenue: $38.7 million."
    )

    doc.add_heading("Digital Transformation", level=2)
    doc.add_paragraph(
        "The Digital Transformation group, managed by Principal Consultant Elena Rodriguez, focused "
        "on process automation and AI integration for manufacturing and logistics clients. Fourteen "
        "RPA implementations were completed, with an average ROI of 340% within 12 months. "
        "Division revenue: $44.9 million."
    )

    # --- Page 4: Client Portfolio ---
    doc.add_page_break()
    doc.add_heading("Client Portfolio Analysis", level=1)
    doc.add_paragraph(
        "Meridian Solutions serves a diversified client base spanning financial services, healthcare, "
        "manufacturing, retail, and government sectors. Client retention rate for FY2025 stood at "
        "94.3%, reflecting strong account management and service delivery."
    )

    client_table = doc.add_table(rows=9, cols=4)
    client_table.style = "Table Grid"
    c_headers = ["Client Name", "Sector", "Contract Value ($M)", "Engagement Type"]
    for i, h_text in enumerate(c_headers):
        cell = client_table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)

    clients = [
        ["Continental Banking Group", "Financial Services", "$8.4", "Cloud Migration"],
        ["Nexus Pharmaceuticals", "Healthcare", "$6.1", "Hybrid Cloud"],
        ["Atlas Manufacturing Corp", "Manufacturing", "$5.7", "Digital Transformation"],
        ["Pinnacle Retail Holdings", "Retail", "$4.2", "Cybersecurity"],
        ["Metro Transit Authority", "Government", "$3.8", "Infrastructure Modernization"],
        ["Horizon Energy Partners", "Energy", "$3.5", "Data Analytics"],
        ["Sapphire Insurance Ltd", "Financial Services", "$3.1", "MDR Services"],
        ["Vertex Logistics Inc", "Logistics", "$2.9", "RPA Implementation"],
    ]
    for r, row_data in enumerate(clients, 1):
        for c, val in enumerate(row_data):
            client_table.cell(r, c).text = val

    doc.add_paragraph("")
    doc.add_paragraph(
        "The top 20 clients represent 61% of total revenue, down from 67% in FY2024, indicating "
        "improved revenue diversification. New client acquisitions totaled 34 in FY2025."
    )

    # --- Page 5: Human Resources ---
    doc.add_page_break()
    doc.add_heading("Human Resources & Talent", level=1)
    doc.add_paragraph(
        "Headcount grew from 842 employees at the start of FY2025 to 1,017 by year-end, including "
        "89 employees who joined through the DataBridge Analytics acquisition. Voluntary turnover "
        "declined to 11.2%, the lowest in five years, attributed to enhanced compensation packages "
        "and the launch of the Meridian Career Accelerator program."
    )
    doc.add_paragraph(
        "Key HR metrics for FY2025:")
    doc.add_paragraph("Total headcount: 1,017", style="List Bullet")
    doc.add_paragraph("New hires (organic): 127", style="List Bullet")
    doc.add_paragraph("Acquisition hires: 89", style="List Bullet")
    doc.add_paragraph("Voluntary turnover rate: 11.2%", style="List Bullet")
    doc.add_paragraph("Average tenure: 4.3 years", style="List Bullet")
    doc.add_paragraph("Employee satisfaction score: 4.2/5.0", style="List Bullet")
    doc.add_paragraph("Training hours per employee: 48", style="List Bullet")

    doc.add_paragraph(
        "The company invested $2.8 million in professional development programs, including "
        "cloud certification sponsorships (AWS, Azure, GCP), leadership workshops, and an internal "
        "mentorship platform connecting junior consultants with senior practitioners."
    )

    # --- Page 6: Technology & Innovation ---
    doc.add_page_break()
    doc.add_heading("Technology & Innovation Initiatives", level=1)
    doc.add_paragraph(
        "Meridian's internal R&D team launched three new proprietary tools in FY2025:"
    )
    doc.add_paragraph(
        "CloudScope — An automated cloud cost optimization platform that analyzes usage patterns "
        "and recommends resource adjustments. Currently deployed across 28 client environments, "
        "achieving an average 23% reduction in monthly cloud spend.",
    )
    doc.add_paragraph(
        "SecureVault — A centralized secrets management solution integrated with CI/CD pipelines. "
        "Adopted by 15 clients within the first quarter of release, reducing credential exposure "
        "incidents by 91%.",
    )
    doc.add_paragraph(
        "FlowEngine — A low-code process automation framework designed for non-technical business "
        "users. Piloted with four enterprise clients in Q4 2025 with full commercial release "
        "planned for Q2 2026.",
    )
    doc.add_paragraph(
        "R&D expenditure for FY2025 totaled $9.6 million, representing 5.1% of revenue. The "
        "innovation pipeline includes seven additional tools in various stages of development."
    )

    # --- Page 7: Strategic Outlook ---
    doc.add_page_break()
    doc.add_heading("Strategic Outlook — FY2026", level=1)
    doc.add_paragraph(
        "The executive leadership team has outlined four strategic priorities for FY2026:"
    )
    doc.add_paragraph(
        "1. Geographic Expansion — Establish a regional office in Singapore to better serve "
        "the growing Asia-Pacific client base, with a target of 15 new APAC engagements by Q3 2026."
    )
    doc.add_paragraph(
        "2. AI-Native Services — Launch a dedicated AI/ML practice offering model development, "
        "MLOps consulting, and AI governance advisory services. Initial team of 25 specialists "
        "to be recruited by Q1 2026."
    )
    doc.add_paragraph(
        "3. Partner Ecosystem — Formalize strategic alliances with three major cloud providers "
        "and two leading cybersecurity vendors to create co-selling opportunities and joint "
        "solution development."
    )
    doc.add_paragraph(
        "4. Sustainability Commitment — Achieve carbon-neutral operations by December 2026, "
        "including transitioning all internal workloads to renewable-energy-powered data centers."
    )
    doc.add_paragraph(
        "Revenue target for FY2026 is $215 million, representing a projected 14.7% growth rate. "
        "Operating margin is expected to reach 18.5% through continued scale efficiencies and "
        "higher-value service mix."
    )

    # --- Page 8: Appendix ---
    doc.add_page_break()
    doc.add_heading("Appendix: Board of Directors & Executive Team", level=1)

    board_table = doc.add_table(rows=7, cols=3)
    board_table.style = "Table Grid"
    b_headers = ["Name", "Title", "Tenure"]
    for i, h_text in enumerate(b_headers):
        cell = board_table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)

    board = [
        ["David Mitchell", "CEO & Chairman", "Since 2011"],
        ["Rachel Torres", "Chief Financial Officer", "Since 2015"],
        ["James Okafor", "Chief Technology Officer", "Since 2018"],
        ["Sarah Chen", "VP, Cloud Infrastructure", "Since 2016"],
        ["Marcus Johnson", "Director, Cybersecurity", "Since 2019"],
        ["Elena Rodriguez", "Principal, Digital Transformation", "Since 2020"],
    ]
    for r, row_data in enumerate(board, 1):
        for c, val in enumerate(row_data):
            board_table.cell(r, c).text = val

    doc.add_paragraph("")
    doc.add_paragraph(
        "This report was prepared by the Office of the CFO in collaboration with divisional "
        "leaders. For questions or clarifications, contact investor.relations@meridiansolutions.com."
    )
    doc.add_paragraph(
        "Meridian Solutions Inc. | 4500 Lakeline Boulevard, Suite 300, Austin, TX 78734 | "
        "www.meridiansolutions.com"
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
