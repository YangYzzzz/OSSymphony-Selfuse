"""
Reward Script: Different first page formatting in Writer document
Task ID: writer_biz_035
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30) - titlePg element enabled (different first page)
  Component 2 (0.25) - First page header is empty / has no content
  Component 3 (0.25) - First page footer is empty / has no content
  Component 4 (0.20) - Default header has company name AND footer has page number field
                        (compound check: only scores if titlePg is also set)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_035'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) < 1:
        print("FAIL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    sect_pr = section._sectPr

    # ---------------------------------------------------------------
    # Component 1: titlePg element is present (0.30 points)
    # This enables "Different First Page" in Writer.
    # INITIAL: titlePg absent. GOLDEN: titlePg present.
    # ---------------------------------------------------------------
    title_pg_present = False
    try:
        title_pg_elem = sect_pr.find(qn('w:titlePg'))
        if title_pg_elem is not None:
            title_pg_present = True
            print(f"PASS: Component 1 — titlePg element found (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 1 — titlePg element not found in sectPr")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: First page header is empty (0.25 points)
    # When titlePg is enabled, a 'first' type headerReference should
    # exist and its content should be empty (no header on first page).
    # INITIAL: No first-page header at all. GOLDEN: Empty first-page header.
    # ---------------------------------------------------------------
    try:
        # Check for first-type headerReference in sectPr
        first_hdr_ref = None
        for child in sect_pr:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'headerReference':
                wtype = child.attrib.get(qn('w:type'), '')
                if wtype == 'first':
                    first_hdr_ref = child
                    break

        if first_hdr_ref is not None and title_pg_present:
            # Access first_page_header via python-docx API
            fp_header = section.first_page_header
            header_text = ''
            if fp_header.paragraphs:
                header_text = ''.join(p.text for p in fp_header.paragraphs).strip()

            if header_text == '':
                print(f"PASS: Component 2 — First page header exists and is empty (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — First page header has content: {repr(header_text)}")
        elif not title_pg_present:
            print(f"FAIL: Component 2 — titlePg not set, so first page header is not differentiated")
        else:
            print(f"FAIL: Component 2 — No first-type headerReference found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: First page footer is empty (0.25 points)
    # Same logic: first-type footerReference should exist and be empty.
    # INITIAL: No first-page footer. GOLDEN: Empty first-page footer.
    # ---------------------------------------------------------------
    try:
        first_ftr_ref = None
        for child in sect_pr:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'footerReference':
                wtype = child.attrib.get(qn('w:type'), '')
                if wtype == 'first':
                    first_ftr_ref = child
                    break

        if first_ftr_ref is not None and title_pg_present:
            fp_footer = section.first_page_footer
            footer_text = ''
            if fp_footer.paragraphs:
                footer_text = ''.join(p.text for p in fp_footer.paragraphs).strip()

            if footer_text == '':
                print(f"PASS: Component 3 — First page footer exists and is empty (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — First page footer has content: {repr(footer_text)}")
        elif not title_pg_present:
            print(f"FAIL: Component 3 — titlePg not set, so first page footer is not differentiated")
        else:
            print(f"FAIL: Component 3 — No first-type footerReference found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Default header has company name AND footer has page
    # number field code — but ONLY if titlePg is enabled (0.20 points)
    # This compound check ensures subsequent pages retain correct
    # header/footer. It cannot score on initial because titlePg is absent.
    # INITIAL: titlePg absent → 0. GOLDEN: titlePg present + header/footer correct → 0.20.
    # ---------------------------------------------------------------
    try:
        if not title_pg_present:
            print(f"FAIL: Component 4 — titlePg not set, compound check skipped")
        else:
            # Check default header for company name
            default_header = section.header
            hdr_text = ''
            if default_header.paragraphs:
                hdr_text = ''.join(p.text for p in default_header.paragraphs).strip()

            has_company = 'Meridian Solutions' in hdr_text

            # Check default footer for PAGE field code
            default_footer = section.footer
            has_page_field = False
            if default_footer._element is not None:
                footer_xml = default_footer._element.xml
                if 'PAGE' in footer_xml or 'page' in footer_xml:
                    has_page_field = True

            if has_company and has_page_field:
                print(f"PASS: Component 4 — Default header has '{hdr_text}' and footer has page number field (0.20 pts)")
                total_score += 0.20
            elif has_company:
                print(f"FAIL: Component 4 — Header OK but footer missing page number field")
            elif has_page_field:
                print(f"FAIL: Component 4 — Footer OK but header missing company name (found: {repr(hdr_text)})")
            else:
                print(f"FAIL: Component 4 — Both header company name and footer page field missing")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
