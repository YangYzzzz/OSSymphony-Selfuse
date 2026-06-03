"""
Reward Script: Split document into two columns from 'Command Reference' onward
Task ID: writer_tech_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Document has at least 2 sections (section break was added)
  Component 2 (0.30): Section break before 'Command Reference' creates a 2-column section
  Component 3 (0.20): The first section (introduction) remains single-column
  Component 4 (0.15): Section break is continuous (not new page)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_050'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_column_count(sectPr):
    """Extract column count from a sectPr element."""
    cols_elem = sectPr.find(qn('w:cols'))
    if cols_elem is not None:
        num = cols_elem.get(qn('w:num'))
        if num is not None:
            return int(num)
    # Default is 1 column
    return 1


def find_section_break_location(doc):
    """
    Find inline section breaks (sectPr inside paragraph pPr).
    Returns list of (para_index, sectPr_element) for each inline break.
    """
    breaks = []
    for i, para in enumerate(doc.paragraphs):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            sect = pPr.find(qn('w:sectPr'))
            if sect is not None:
                breaks.append((i, sect))
    return breaks


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)

    # Component 1: Document has at least 2 sections (0.35 points)
    # Initial doc has only 1 section; golden has 2 (section break added)
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 - Document has {num_sections} sections (>=2) (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 - Document has only {num_sections} section(s), expected >=2")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: A section with 2 columns exists, and 'Command Reference' heading
    # appears in or after that section (0.30 points)
    try:
        two_col_section_idx = None
        for i, sec in enumerate(doc.sections):
            col_count = get_column_count(sec._sectPr)
            if col_count >= 2:
                two_col_section_idx = i
                break

        if two_col_section_idx is not None:
            # Check that 'Command Reference' is in the two-column section
            # The inline sectPr before 'Command Reference' defines the END of the
            # first section. So 'Command Reference' should be AFTER the break.
            inline_breaks = find_section_break_location(doc)
            cmd_ref_para_idx = None
            for i, para in enumerate(doc.paragraphs):
                if 'Command Reference' in para.text:
                    cmd_ref_para_idx = i
                    break

            if cmd_ref_para_idx is not None:
                # The section break should be at or before the Command Reference para
                # In docx, inline sectPr in a paragraph means that paragraph is the
                # LAST paragraph of the section. So the break para index should be
                # less than cmd_ref_para_idx.
                break_before_cmd_ref = any(
                    bi < cmd_ref_para_idx for bi, _ in inline_breaks
                )
                if break_before_cmd_ref:
                    print(f"PASS: Component 2 - 2-column section found at index {two_col_section_idx}, "
                          f"'Command Reference' at para {cmd_ref_para_idx} is after section break (0.30 pts)")
                    total_score += 0.30
                else:
                    print(f"FAIL: Component 2 - 2-column section exists but section break is not before "
                          f"'Command Reference' (para {cmd_ref_para_idx})")
            else:
                print(f"FAIL: Component 2 - 'Command Reference' heading not found in document")
        else:
            print(f"FAIL: Component 2 - No section with 2+ columns found")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: First section (introduction) remains single-column (0.20 points)
    # This only matters if there are multiple sections - anchored to the task change
    try:
        if num_sections >= 2:
            first_sec_cols = get_column_count(doc.sections[0]._sectPr)
            if first_sec_cols == 1:
                print(f"PASS: Component 3 - First section is single-column ({first_sec_cols} col) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 - First section has {first_sec_cols} columns, expected 1")
        else:
            print(f"FAIL: Component 3 - Only 1 section exists; cannot verify intro is single-column "
                  f"while Command Reference is two-column")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: The section break is continuous (not new page) (0.15 points)
    # A continuous break keeps content flowing on the same page
    try:
        if num_sections >= 2:
            # Check the second section's start_type
            from docx.enum.section import WD_SECTION_START
            second_sec_start = doc.sections[1].start_type
            if second_sec_start == WD_SECTION_START.CONTINUOUS:
                print(f"PASS: Component 4 - Section break is CONTINUOUS (0.15 pts)")
                total_score += 0.15
            else:
                # Also check inline sectPr for the break type
                inline_breaks = find_section_break_location(doc)
                continuous_found = any(
                    sect_elem.find(qn('w:type')) is not None
                    and sect_elem.find(qn('w:type')).get(qn('w:val'), '') == 'continuous'
                    for _, sect_elem in inline_breaks
                )
                if continuous_found:
                    print(f"PASS: Component 4 - Inline section break is continuous (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 4 - Section break type is {second_sec_start}, expected CONTINUOUS")
        else:
            print(f"FAIL: Component 4 - No section break to evaluate")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
