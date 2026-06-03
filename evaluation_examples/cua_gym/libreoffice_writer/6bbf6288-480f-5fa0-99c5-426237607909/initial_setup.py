"""
Initial Setup: Single-column tech document with introduction and command reference
Task ID: writer_tech_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ---- Introduction ----
    title = doc.add_heading('NetForge CLI Toolkit', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('Introduction', level=1)

    p1 = doc.add_paragraph(
        'NetForge is a comprehensive command-line toolkit designed for network engineers '
        'and system administrators who manage large-scale infrastructure deployments. '
        'Built on a modular plugin architecture, NetForge provides a unified interface '
        'for provisioning, monitoring, and troubleshooting network devices across '
        'heterogeneous environments.'
    )

    p2 = doc.add_paragraph(
        'Originally developed at Meridian Systems in 2019, NetForge has grown from a '
        'simple SSH wrapper into a full-featured automation platform supporting over '
        '200 device types from vendors including Cisco, Juniper, Arista, and Palo Alto '
        'Networks. The toolkit integrates seamlessly with existing CI/CD pipelines and '
        'configuration management tools such as Ansible and Terraform.'
    )

    p3 = doc.add_paragraph(
        'This document serves as the primary reference for NetForge version 4.2.1, '
        'released on March 15, 2025. It covers installation prerequisites, core '
        'commands, configuration options, and troubleshooting procedures. For the '
        'latest updates and community plugins, visit the project repository at '
        'https://github.com/meridian-systems/netforge.'
    )

    doc.add_heading('System Requirements', level=2)

    req_text = doc.add_paragraph(
        'Before installing NetForge, ensure your system meets the following minimum '
        'requirements:'
    )

    doc.add_paragraph('Python 3.9 or later (3.11+ recommended)', style='List Bullet')
    doc.add_paragraph('OpenSSH 8.0 or later with multiplexing support', style='List Bullet')
    doc.add_paragraph('At least 512 MB of available RAM for the daemon process', style='List Bullet')
    doc.add_paragraph('Network access to target devices on ports 22 (SSH) and 830 (NETCONF)', style='List Bullet')
    doc.add_paragraph('libffi and libssl development headers for cryptographic operations', style='List Bullet')

    p4 = doc.add_paragraph(
        'NetForge supports Linux (Ubuntu 20.04+, RHEL 8+, Debian 11+), macOS 12+, '
        'and Windows 10/11 via WSL2. Native Windows support is experimental and not '
        'recommended for production environments.'
    )

    # ---- Command Reference ----
    doc.add_heading('Command Reference', level=1)

    p5 = doc.add_paragraph(
        'The following commands form the core of the NetForge toolkit. Each command '
        'supports the --help flag for detailed usage information and the --verbose '
        'flag for debug output.'
    )

    # Command: nf init
    doc.add_heading('nf init', level=2)
    doc.add_paragraph(
        'Initializes a new NetForge project in the current directory. Creates the '
        'configuration file (netforge.yaml), inventory directory, and plugin cache.'
    )
    p_usage1 = doc.add_paragraph()
    run_u1 = p_usage1.add_run('Usage: ')
    run_u1.bold = True
    p_usage1.add_run('nf init [--template <name>] [--inventory <path>]')

    p_opt1 = doc.add_paragraph()
    run_o1 = p_opt1.add_run('Options: ')
    run_o1.bold = True
    p_opt1.add_run('--template specifies a starter template (default: minimal). '
                   '--inventory sets a custom inventory path.')

    # Command: nf discover
    doc.add_heading('nf discover', level=2)
    doc.add_paragraph(
        'Scans the specified network range and automatically detects devices, their '
        'types, and firmware versions. Results are written to the inventory file.'
    )
    p_usage2 = doc.add_paragraph()
    run_u2 = p_usage2.add_run('Usage: ')
    run_u2.bold = True
    p_usage2.add_run('nf discover <subnet> [--timeout <seconds>] [--protocol ssh|snmp]')

    p_opt2 = doc.add_paragraph()
    run_o2 = p_opt2.add_run('Options: ')
    run_o2.bold = True
    p_opt2.add_run('--timeout sets the per-device connection timeout (default: 30s). '
                   '--protocol selects the discovery method.')

    # Command: nf deploy
    doc.add_heading('nf deploy', level=2)
    doc.add_paragraph(
        'Pushes configuration changes to one or more devices. Supports dry-run mode, '
        'automatic rollback on failure, and parallel execution across device groups.'
    )
    p_usage3 = doc.add_paragraph()
    run_u3 = p_usage3.add_run('Usage: ')
    run_u3.bold = True
    p_usage3.add_run('nf deploy [--target <group>] [--dry-run] [--parallel <n>]')

    p_opt3 = doc.add_paragraph()
    run_o3 = p_opt3.add_run('Options: ')
    run_o3.bold = True
    p_opt3.add_run('--target restricts deployment to a named group. --dry-run shows '
                   'changes without applying. --parallel sets concurrency (default: 5).')

    # Command: nf monitor
    doc.add_heading('nf monitor', level=2)
    doc.add_paragraph(
        'Starts real-time monitoring of device health metrics including CPU utilization, '
        'memory usage, interface errors, and BGP session state. Data is streamed to the '
        'console or forwarded to a configured metrics backend (Prometheus, InfluxDB).'
    )
    p_usage4 = doc.add_paragraph()
    run_u4 = p_usage4.add_run('Usage: ')
    run_u4.bold = True
    p_usage4.add_run('nf monitor [--interval <seconds>] [--export <backend>]')

    p_opt4 = doc.add_paragraph()
    run_o4 = p_opt4.add_run('Options: ')
    run_o4.bold = True
    p_opt4.add_run('--interval sets the polling frequency (default: 60s). --export '
                   'sends data to Prometheus or InfluxDB.')

    # Command: nf backup
    doc.add_heading('nf backup', level=2)
    doc.add_paragraph(
        'Creates a snapshot of running configurations from all devices in the inventory. '
        'Backups are stored locally with timestamps and can be pushed to a Git repository '
        'for version control.'
    )
    p_usage5 = doc.add_paragraph()
    run_u5 = p_usage5.add_run('Usage: ')
    run_u5.bold = True
    p_usage5.add_run('nf backup [--output <dir>] [--git-push] [--diff-only]')

    p_opt5 = doc.add_paragraph()
    run_o5 = p_opt5.add_run('Options: ')
    run_o5.bold = True
    p_opt5.add_run('--output sets the backup directory (default: ./backups). --git-push '
                   'commits and pushes changes. --diff-only stores only changed configs.')

    # Command: nf audit
    doc.add_heading('nf audit', level=2)
    doc.add_paragraph(
        'Runs compliance checks against device configurations using predefined or custom '
        'rule sets. Generates a report highlighting violations, risk scores, and '
        'recommended remediation steps.'
    )
    p_usage6 = doc.add_paragraph()
    run_u6 = p_usage6.add_run('Usage: ')
    run_u6.bold = True
    p_usage6.add_run('nf audit [--ruleset <name>] [--format json|html|pdf]')

    p_opt6 = doc.add_paragraph()
    run_o6 = p_opt6.add_run('Options: ')
    run_o6.bold = True
    p_opt6.add_run('--ruleset selects the compliance standard (cis, nist, custom). '
                   '--format sets the output report format (default: html).')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
