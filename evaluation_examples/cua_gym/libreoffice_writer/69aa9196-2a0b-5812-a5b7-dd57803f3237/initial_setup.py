"""
Initial Setup: Memo document with mixed formatting for copy-paste-as-unformatted task
Task ID: writer_edit_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user/Desktop'  # VM path — file goes on Desktop as per task context
TASK_ID = 'formatted_memo'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Paragraph 1: Bold title "Internal Memo"
    para_title = doc.add_paragraph()
    run_title = para_title.add_run("Internal Memo")
    run_title.bold = True
    run_title.font.size = Pt(14)

    # Paragraph 2: Italic date "Date: March 3, 2025"
    para_date = doc.add_paragraph()
    run_date = para_date.add_run("Date: March 3, 2025")
    run_date.italic = True

    # Paragraph 3: Subject line (no special formatting)
    para_subject = doc.add_paragraph()
    run_subject = para_subject.add_run("Subject: Office Relocation")

    # Paragraph 4: Body text with underlined key phrases
    para_body = doc.add_paragraph()
    run_body1 = para_body.add_run("We are pleased to announce that our office will relocate to ")
    run_body_addr = para_body.add_run("123 Main Street")
    run_body_addr.underline = True
    run_body2 = para_body.add_run(" effective ")
    run_body_date = para_body.add_run("April 1, 2025")
    run_body_date.underline = True
    run_body3 = para_body.add_run(". Please contact HR for details.")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
