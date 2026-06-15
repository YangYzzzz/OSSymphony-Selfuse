"""
Reward Script: Conference agenda document verification
Task ID: writer_wf_017
Domain: libreoffice_writer
Scoring:
  Component 1: Title 'Annual Tech Summit 2025' present (0.15)
  Component 2: Date/venue line below title (0.10)
  Component 3: Table with 4 columns and 9 rows (0.20)
  Component 4: Table header row has Time/Session/Speaker/Room (0.10)
  Component 5: 8 sessions spanning 9AM-5PM (0.15)
  Component 6: Header row bold formatting (0.10)
  Component 7: Header row blue shading (0.10)
  Component 8: Footer 'Subject to change' (0.10)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_017'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify conference agenda document with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title 'Annual Tech Summit 2025' (0.15 points)
    try:
        title_found = False
        for para in doc.paragraphs:
            if 'annual tech summit 2025' in para.text.lower():
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 — Title 'Annual Tech Summit 2025' found (0.15 pts)")
            total_score += 0.15
        else:
            all_text = [p.text for p in doc.paragraphs]
            print(f"FAIL: Component 1 — Title not found. Paragraphs: {all_text[:5]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Date/venue info below title (0.10 points)
    try:
        date_venue_found = False
        for para in doc.paragraphs:
            text_lower = para.text.lower()
            # Check for date-like content AND venue/location reference
            has_date = bool(re.search(r'\d{4}', para.text)) or bool(re.search(r'(january|february|march|april|may|june|july|august|september|october|november|december)', text_lower))
            has_venue = any(word in text_lower for word in ['center', 'hall', 'convention', 'venue', 'hotel', 'san francisco', 'location'])
            if has_date and has_venue:
                date_venue_found = True
                break
        if date_venue_found:
            print(f"PASS: Component 2 — Date and venue info found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — No paragraph with both date and venue info")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table with 4 columns and >= 9 rows (0.20 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_cols == 4 and num_rows >= 9:
                print(f"PASS: Component 3 — Table has {num_rows} rows x {num_cols} cols (0.20 pts)")
                total_score += 0.20
            elif num_cols == 4:
                # Partial: right columns but fewer rows
                partial = 0.10
                print(f"PARTIAL: Component 3 — Table has {num_rows} rows x {num_cols} cols (expected >= 9 rows) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Table has {num_rows} rows x {num_cols} cols (expected 4 cols, >= 9 rows)")
        else:
            print(f"FAIL: Component 3 — No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Header row has Time, Session, Speaker, Room (0.10 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            expected = ['time', 'session', 'speaker', 'room']
            matches = sum(1 for exp in expected if any(exp in h for h in header_cells))
            if matches == 4:
                print(f"PASS: Component 4 — Header row has all 4 expected columns (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Header row: {header_cells}, matched {matches}/4 expected columns")
        else:
            print(f"FAIL: Component 4 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 8 sessions spanning 9AM-5PM including required types (0.15 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            # Check data rows (skip header)
            data_rows = []
            for ri in range(1, len(table.rows)):
                row_text = [cell.text.strip().lower() for cell in table.rows[ri].cells]
                data_rows.append(row_text)

            session_count = len(data_rows)

            # Check for required session types
            all_text = ' '.join(' '.join(row) for row in data_rows)
            has_registration = 'registration' in all_text or 'welcome' in all_text
            has_keynote = 'keynote' in all_text
            has_lunch = 'lunch' in all_text
            has_closing = 'closing' in all_text or 'closing remarks' in all_text
            has_breakout = all_text.count('breakout') >= 4

            required_checks = [has_registration, has_keynote, has_lunch, has_closing, has_breakout]
            passed = sum(required_checks)

            if session_count >= 8 and passed == 5:
                print(f"PASS: Component 5 — {session_count} sessions with all required types (0.15 pts)")
                total_score += 0.15
            elif session_count >= 6 and passed >= 3:
                partial = 0.08
                print(f"PARTIAL: Component 5 — {session_count} sessions, {passed}/5 required types ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — {session_count} sessions, required types: reg={has_registration} key={has_keynote} lunch={has_lunch} close={has_closing} breakout={has_breakout}")
        else:
            print(f"FAIL: Component 5 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Header row bold formatting (0.10 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            bold_count = 0
            total_header_runs = 0
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            total_header_runs += 1
                            if run.font.bold:
                                bold_count += 1

            if total_header_runs > 0 and bold_count >= total_header_runs:
                print(f"PASS: Component 6 — Header row is bold ({bold_count}/{total_header_runs} runs) (0.10 pts)")
                total_score += 0.10
            elif total_header_runs > 0 and bold_count > 0:
                partial = 0.05
                print(f"PARTIAL: Component 6 — Some header runs bold ({bold_count}/{total_header_runs}) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — Header row not bold ({bold_count}/{total_header_runs} runs)")
        else:
            print(f"FAIL: Component 6 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Header row blue shading (0.10 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            blue_shaded = 0
            for cell in table.rows[0].cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill:
                            fill = fill.upper()
                            # Check if the fill is a blue-ish color
                            # Blue colors typically have high B channel relative to R and G
                            try:
                                r_val = int(fill[0:2], 16)
                                g_val = int(fill[2:4], 16)
                                b_val = int(fill[4:6], 16)
                                # Blue dominant: B > R and B > G, or known blue hex
                                if b_val > r_val and b_val > 100:
                                    blue_shaded += 1
                            except (ValueError, IndexError):
                                pass

            if blue_shaded >= 4:
                print(f"PASS: Component 7 — All 4 header cells have blue shading (0.10 pts)")
                total_score += 0.10
            elif blue_shaded > 0:
                partial = 0.05
                print(f"PARTIAL: Component 7 — {blue_shaded}/4 header cells have blue shading ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 — No blue shading found on header row cells")
        else:
            print(f"FAIL: Component 7 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: Footer contains 'Subject to change' (0.10 points)
    try:
        footer_found = False
        for section in doc.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                footer_text = ' '.join(fp.text for fp in footer.paragraphs).strip().lower()
                if 'subject to change' in footer_text:
                    footer_found = True
                    break
        if footer_found:
            print(f"PASS: Component 8 — Footer contains 'Subject to change' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 8 — Footer does not contain 'Subject to change'")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
