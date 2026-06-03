"""
Reward Script: Set up paragraph styles for marketing document
Task ID: writer_mktg_022
Domain: libreoffice_writer
Scoring:
  Component 1: 'Marketing Body' style exists with correct properties (0.25 pts)
               - Calibri 12pt, line spacing 1.15, space after 6pt
  Component 2: 'Marketing Heading' style exists with correct properties (0.25 pts)
               - Calibri 16pt bold, dark blue #0D47A1, space before 12pt, space after 6pt
  Component 3: All heading paragraphs assigned 'Marketing Heading' style (0.25 pts)
  Component 4: All body paragraphs assigned 'Marketing Body' style (0.25 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_022'
FILE_PATH = '/home/user/Desktop/marketing_style_guide.docx'

# Tolerance for floating point comparisons (EMU units)
EMU_TOLERANCE = 100  # small tolerance for rounding differences

def emu_to_pt(emu):
    """Convert EMU to points."""
    if emu is None:
        return None
    return emu / 12700.0

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Build a dict of style name -> style object for convenience
    style_map = {s.name: s for s in doc.styles}

    # -----------------------------------------------------------------------
    # Component 1: 'Marketing Body' style exists with correct properties (0.25 pts)
    # Properties: Calibri, 12pt, line spacing 1.15, space after 6pt
    # -----------------------------------------------------------------------
    try:
        mb_exists = 'Marketing Body' in style_map
        if not mb_exists:
            print("FAIL: Component 1 — 'Marketing Body' style does not exist")
        if mb_exists:
            mb_style = style_map['Marketing Body']
            mb_pf = mb_style.paragraph_format
            mb_font = mb_style.font

            issues = []

            # Check font name == Calibri
            if mb_font.name != 'Calibri':
                issues.append(f"font.name={mb_font.name!r} (expected 'Calibri')")

            # Check font size == 12pt (152400 EMU)
            if mb_font.size is None or abs(mb_font.size - Pt(12)) > EMU_TOLERANCE:
                issues.append(f"font.size={emu_to_pt(mb_font.size)}pt (expected 12pt)")

            # Check line spacing == 1.15
            if mb_pf.line_spacing is None or abs(float(mb_pf.line_spacing) - 1.15) > 0.02:
                issues.append(f"line_spacing={mb_pf.line_spacing} (expected 1.15)")

            # Check space after == 6pt (76200 EMU)
            if mb_pf.space_after is None or abs(mb_pf.space_after - Pt(6)) > EMU_TOLERANCE:
                issues.append(f"space_after={emu_to_pt(mb_pf.space_after)}pt (expected 6pt)")

            if issues:
                print(f"FAIL: Component 1 — 'Marketing Body' style issues: {'; '.join(issues)}")
            if not issues:
                print(f"PASS: Component 1 — 'Marketing Body' style has correct properties: "
                      f"Calibri 12pt, line_spacing=1.15, space_after=6pt (0.25 pts)")
                total_score += 0.25
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: 'Marketing Heading' style exists with correct properties (0.25 pts)
    # Properties: Calibri 16pt bold, dark blue #0D47A1, space before 12pt, space after 6pt
    # -----------------------------------------------------------------------
    try:
        mh_exists = 'Marketing Heading' in style_map
        if not mh_exists:
            print("FAIL: Component 2 — 'Marketing Heading' style does not exist")
        if mh_exists:
            mh_style = style_map['Marketing Heading']
            mh_pf = mh_style.paragraph_format
            mh_font = mh_style.font

            issues = []

            # Check font name == Calibri
            if mh_font.name != 'Calibri':
                issues.append(f"font.name={mh_font.name!r} (expected 'Calibri')")

            # Check font size == 16pt (203200 EMU)
            if mh_font.size is None or abs(mh_font.size - Pt(16)) > EMU_TOLERANCE:
                issues.append(f"font.size={emu_to_pt(mh_font.size)}pt (expected 16pt)")

            # Check bold == True
            if not mh_font.bold:
                issues.append(f"font.bold={mh_font.bold} (expected True)")

            # Check color == dark blue #0D47A1
            # RGBColor in python-docx is compared as a hex string (e.g., '0D47A1')
            try:
                rgb = mh_font.color.rgb
                actual_hex = str(rgb).upper()
                expected_hex = '0D47A1'
                if actual_hex != expected_hex:
                    issues.append(f"font.color={actual_hex!r} (expected '0D47A1')")
            except Exception as ce:
                issues.append(f"font.color could not be read: {ce}")

            # Check space before == 12pt (152400 EMU)
            if mh_pf.space_before is None or abs(mh_pf.space_before - Pt(12)) > EMU_TOLERANCE:
                issues.append(f"space_before={emu_to_pt(mh_pf.space_before)}pt (expected 12pt)")

            # Check space after == 6pt (76200 EMU)
            if mh_pf.space_after is None or abs(mh_pf.space_after - Pt(6)) > EMU_TOLERANCE:
                issues.append(f"space_after={emu_to_pt(mh_pf.space_after)}pt (expected 6pt)")

            if issues:
                print(f"FAIL: Component 2 — 'Marketing Heading' style issues: {'; '.join(issues)}")
            if not issues:
                print(f"PASS: Component 2 — 'Marketing Heading' style has correct properties: "
                      f"Calibri 16pt bold, color=#0D47A1, space_before=12pt, space_after=6pt (0.25 pts)")
                total_score += 0.25
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: All heading paragraphs use 'Marketing Heading' style (0.25 pts)
    # The document originally had bold+large-font headings with 'Normal' style.
    # In golden_env, these must be assigned 'Marketing Heading'.
    # We verify: at least one paragraph uses 'Marketing Heading', and all non-empty
    # paragraphs use either 'Marketing Heading' or 'Marketing Body'.
    # -----------------------------------------------------------------------
    try:
        mh_style_present = 'Marketing Heading' in style_map
        if not mh_style_present:
            print("FAIL: Component 3 — 'Marketing Heading' style does not exist (skipped)")
        if mh_style_present:
            heading_paras = [p for p in doc.paragraphs if p.style.name == 'Marketing Heading']
            non_target_styles = set()
            for p in doc.paragraphs:
                if p.text.strip() and p.style.name not in ('Marketing Heading', 'Marketing Body'):
                    non_target_styles.add(p.style.name)

            has_heading_paras = len(heading_paras) > 0
            no_unexpected_styles = len(non_target_styles) == 0

            if not has_heading_paras:
                print("FAIL: Component 3 — No paragraphs use 'Marketing Heading' style")
            if has_heading_paras and not no_unexpected_styles:
                print(f"FAIL: Component 3 — Some non-empty paragraphs use unexpected styles: {non_target_styles}")
            if has_heading_paras and no_unexpected_styles:
                print(f"PASS: Component 3 — {len(heading_paras)} paragraph(s) assigned 'Marketing Heading' style, "
                      f"no unexpected styles found (0.25 pts)")
                total_score += 0.25
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: All body paragraphs use 'Marketing Body' style (0.25 pts)
    # We verify: at least some paragraphs use 'Marketing Body', and together with
    # Component 3 all non-empty paragraphs are covered by Marketing styles.
    # -----------------------------------------------------------------------
    try:
        mb_style_present = 'Marketing Body' in style_map
        if not mb_style_present:
            print("FAIL: Component 4 — 'Marketing Body' style does not exist (skipped)")
        if mb_style_present:
            body_paras = [p for p in doc.paragraphs if p.style.name == 'Marketing Body']
            total_non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
            heading_paras_count = sum(1 for p in doc.paragraphs if p.style.name == 'Marketing Heading')
            body_paras_count = len(body_paras)

            # All non-empty paras should be covered by Marketing Heading + Marketing Body
            covered = heading_paras_count + body_paras_count
            has_body_paras = body_paras_count > 0
            all_covered = covered >= total_non_empty

            if not has_body_paras:
                print("FAIL: Component 4 — No paragraphs use 'Marketing Body' style")
            if has_body_paras and not all_covered:
                print(f"FAIL: Component 4 — {total_non_empty - covered} non-empty paragraphs not covered "
                      f"by Marketing styles (total_non_empty={total_non_empty}, covered={covered})")
            if has_body_paras and all_covered:
                print(f"PASS: Component 4 — {body_paras_count} paragraph(s) assigned 'Marketing Body' style, "
                      f"all {total_non_empty} non-empty paragraphs covered by Marketing styles (0.25 pts)")
                total_score += 0.25
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
