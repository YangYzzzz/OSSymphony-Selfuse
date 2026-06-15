"""
Initial Setup: Marketing Style Guide Document
Task ID: writer_mktg_022
Domain: libreoffice_writer

Creates a marketing style guide document with:
- Headings manually formatted as 16pt bold (NOT using paragraph styles)
- Body text 12pt with inconsistent fonts (mix of Arial and Calibri)
- NO custom paragraph styles defined
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_022'
OUTPUT = f'{WORKDIR}/marketing_style_guide.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_manual_heading(doc, text, font_name="Calibri"):
    """Add a heading paragraph with manual formatting (no paragraph style)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = font_name
    return para


def add_body_paragraph(doc, text, font_name="Arial"):
    """Add a body paragraph with manual formatting, inconsistent font."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = font_name
    return para


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default styles interference — just use Normal style default
    # Page 1: Company Overview section
    add_manual_heading(doc, "Company Overview", font_name="Calibri")
    add_body_paragraph(doc, (
        "Apex Marketing Solutions was founded in 2015 with a singular vision: to transform how "
        "businesses connect with their customers. Our team of over 120 marketing professionals "
        "brings decades of combined experience across digital, print, and experiential marketing."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "We serve clients across North America, Europe, and Asia-Pacific, delivering integrated "
        "marketing campaigns that drive measurable results. Our award-winning creative team has "
        "produced campaigns for Fortune 500 companies and high-growth startups alike."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Our core values of innovation, integrity, and impact guide every project we undertake. "
        "We believe that great marketing is built on deep consumer insight, compelling storytelling, "
        "and relentless optimization."
    ), font_name="Arial")

    # Page 1-2: Brand Identity section
    add_manual_heading(doc, "Brand Identity Guidelines", font_name="Arial")
    add_body_paragraph(doc, (
        "Consistent brand identity is the foundation of effective marketing. Every touchpoint — "
        "from business cards to billboard ads — must reflect our brand values and visual language. "
        "This section outlines the core elements of our brand identity system."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Our primary logo should always appear in its approved color variants: full-color on white "
        "backgrounds, white on dark backgrounds, and single-color black on specialty applications. "
        "Never stretch, rotate, or add effects to the logo without explicit approval."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "Clear space requirements must be maintained around the logo at all times. The minimum clear "
        "space equals the height of the capital 'A' in our wordmark. This ensures visual breathing "
        "room and prevents the logo from appearing crowded."
    ), font_name="Calibri")

    # Page 2: Color Palette
    add_manual_heading(doc, "Color Palette", font_name="Calibri")
    add_body_paragraph(doc, (
        "Our brand color palette has been carefully selected to convey professionalism, energy, "
        "and approachability. The primary palette consists of Apex Blue (#1A3C6E), Vibrant Orange "
        "(#E85B26), and Clean White (#FFFFFF)."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "Secondary colors may be used for supporting graphics and infographics: Light Blue (#5B9BD5), "
        "Warm Gray (#8C8C8C), and Dark Charcoal (#2D2D2D). These colors should never overpower the "
        "primary palette but serve to add depth and variety to layouts."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Digital applications should use the RGB values specified in this document. Print applications "
        "must use the CMYK breakdowns to ensure accurate reproduction across different printing "
        "technologies and substrates."
    ), font_name="Arial")

    # Page 2-3: Typography
    add_manual_heading(doc, "Typography Standards", font_name="Arial")
    add_body_paragraph(doc, (
        "Typography plays a critical role in communicating our brand personality. Our typeface system "
        "is built around Calibri for primary communications and Arial as the system fallback. This "
        "combination ensures readability across all media while maintaining brand consistency."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Headline typefaces should be used at sizes 24pt and above. For sub-headings, use 16-20pt "
        "with appropriate weight variations. Body copy should consistently appear at 10-12pt "
        "depending on the medium and column width."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "Line spacing and paragraph spacing are critical for readability. Marketing materials should "
        "use 1.15 line spacing as the standard, with 6-8pt space after paragraphs. This creates "
        "comfortable reading rhythm without wasting space on the page."
    ), font_name="Calibri")

    # Page 3: Digital Marketing
    add_manual_heading(doc, "Digital Marketing Channels", font_name="Calibri")
    add_body_paragraph(doc, (
        "Our digital marketing strategy spans six primary channels: search engine optimization, "
        "paid search advertising, social media marketing, email campaigns, content marketing, "
        "and display advertising. Each channel requires tailored content and messaging approaches."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "Social media guidelines require all posts to include a visual element, use approved hashtags, "
        "and maintain the brand voice. LinkedIn content should emphasize thought leadership and "
        "professional insights. Instagram and Facebook focus on lifestyle imagery and community engagement."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Email marketing templates must follow the approved layout grid with a maximum width of 600px. "
        "Subject lines should be 40-60 characters. Preview text must complement the subject line and "
        "encourage opens. All emails must include an unsubscribe link and physical mailing address."
    ), font_name="Arial")

    # Page 3-4: Content Strategy
    add_manual_heading(doc, "Content Strategy Framework", font_name="Arial")
    add_body_paragraph(doc, (
        "Our content strategy is built on the PESO model: Paid, Earned, Shared, and Owned media. "
        "Every piece of content should serve at least one of three objectives: build awareness, "
        "nurture prospects, or convert leads into customers."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Content calendars are maintained on a quarterly basis with monthly reviews. Campaign themes "
        "are aligned with business priorities, seasonal trends, and industry events. Content should "
        "be repurposed across channels wherever possible to maximize production efficiency."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "All published content must pass through our three-stage review process: editorial review "
        "for accuracy and brand voice, legal review for compliance and risk, and final approval "
        "from the marketing director. Turnaround times vary by content type."
    ), font_name="Calibri")

    # Page 4-5: Campaign Measurement
    add_manual_heading(doc, "Campaign Measurement & Analytics", font_name="Calibri")
    add_body_paragraph(doc, (
        "Measurement is integral to our marketing operations. Every campaign must have defined KPIs "
        "established before launch, tracked in real-time during execution, and analyzed post-campaign "
        "for optimization insights."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "Primary metrics we track include: reach, impressions, click-through rate, conversion rate, "
        "cost per acquisition, return on ad spend, brand sentiment score, and net promoter score. "
        "Secondary metrics provide supporting context for primary KPI performance."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Monthly performance reports are distributed to all stakeholders by the 5th of each month. "
        "Quarterly business reviews include competitive benchmarking analysis and strategic "
        "recommendations for the following quarter."
    ), font_name="Arial")

    # Page 5: Approval Workflows
    add_manual_heading(doc, "Approval Workflows & Compliance", font_name="Arial")
    add_body_paragraph(doc, (
        "All marketing materials must follow our approval workflow before publication or distribution. "
        "The workflow ensures brand consistency, legal compliance, and stakeholder alignment. Rushing "
        "approvals without completing all required reviews is strictly prohibited."
    ), font_name="Calibri")
    add_body_paragraph(doc, (
        "Creative briefs must be approved by the account manager and creative director before any "
        "production work begins. The brief should include target audience profile, key message hierarchy, "
        "mandatory elements, and production specifications."
    ), font_name="Arial")
    add_body_paragraph(doc, (
        "Post-campaign documentation including final assets, performance reports, and lessons learned "
        "must be uploaded to our shared drive within 30 days of campaign completion. This builds our "
        "institutional knowledge base and enables future campaign optimization."
    ), font_name="Calibri")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open with LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
