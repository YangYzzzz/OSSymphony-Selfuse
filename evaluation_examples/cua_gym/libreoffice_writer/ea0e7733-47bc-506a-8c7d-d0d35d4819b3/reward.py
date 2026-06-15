"""
Reward Script: Add page numbers in footer for book-style layout
Task ID: writer_page_068
Domain: libreoffice_writer
Scoring:
  Component 1: evenAndOddHeaders enabled in document settings (0.2 pts)
  Component 2: Even-page footer contains PAGE field code (0.25 pts)
  Component 3: Even-page footer is left-aligned (0.25 pts)
  Component 4: Odd-page footer contains PAGE field code (0.15 pts)
  Component 5: Odd-page footer is right-aligned (0.15 pts)
  Total: 1.0
"""

import os

# python-docx for .docx verification
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import lxml.etree as etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_068'
FILE_NAME = 'reference_manual.docx'


def has_page_field(paragraphs):
    """Check if any paragraph in the list contains a PAGE field code (instrText with PAGE)."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for para in paragraphs:
        instrs = para._element.findall('.//w:instrText', ns)
        for instr in instrs:
            if instr.text and 'PAGE' in instr.text.upper():
                return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Task: Add page numbers in the footer:
      - Left-aligned on even pages
      - Right-aligned on odd pages
      - Different content for even/odd pages (not 'same content on both sides')
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file {}: {}".format(file_path, e))
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) == 0:
        print("CRITICAL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: evenAndOddHeaders enabled in document settings (0.2 pts)
    # This flag enables the 'different content on left and right pages' feature.
    # It must be present in document-level settings (w:settings).
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        settings_part = doc.part.part_related_by(RT.SETTINGS)
        settings_xml = etree.tostring(settings_part._element).decode()
        if 'evenAndOddHeaders' in settings_xml:
            print("PASS: Component 1 -- evenAndOddHeaders is enabled in document settings (0.2 pts)")
            total_score += 0.2
        else:
            print("FAIL: Component 1 -- evenAndOddHeaders NOT found in document settings (expected enabled for book-style layout)")
    except Exception as e:
        print("ERROR: Component 1 -- Could not check document settings: {}".format(e))

    # Component 2: Even-page footer contains a PAGE field code (0.25 pts)
    # The even-page footer should contain a PAGE field for the page number.
    try:
        even_footer = section.even_page_footer
        even_has_page = has_page_field(even_footer.paragraphs)
        if even_has_page:
            print("PASS: Component 2 -- Even-page footer contains PAGE field code (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 2 -- Even-page footer does NOT contain a PAGE field code")
    except Exception as e:
        print("ERROR: Component 2 -- Could not check even-page footer field: {}".format(e))

    # Component 3: Even-page footer is left-aligned (0.25 pts)
    # The task requires page numbers left-aligned on even pages.
    try:
        even_footer = section.even_page_footer
        even_alignments = [
            para.paragraph_format.alignment
            for para in even_footer.paragraphs
        ]
        even_has_left = any(a == WD_PARAGRAPH_ALIGNMENT.LEFT for a in even_alignments)
        if even_has_left:
            print("PASS: Component 3 -- Even-page footer paragraph is LEFT-aligned (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 3 -- Even-page footer is NOT left-aligned; found alignments: {}".format(even_alignments))
    except Exception as e:
        print("ERROR: Component 3 -- Could not check even-page footer alignment: {}".format(e))

    # Component 4: Odd-page footer (default footer) contains a PAGE field code (0.15 pts)
    # The default/odd-page footer should contain a PAGE field for the page number.
    try:
        odd_footer = section.footer
        odd_has_page = has_page_field(odd_footer.paragraphs)
        if odd_has_page:
            print("PASS: Component 4 -- Odd-page footer contains PAGE field code (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 4 -- Odd-page footer does NOT contain a PAGE field code")
    except Exception as e:
        print("ERROR: Component 4 -- Could not check odd-page footer field: {}".format(e))

    # Component 5: Odd-page footer (default footer) is right-aligned (0.15 pts)
    # The task requires page numbers right-aligned on odd pages.
    try:
        odd_footer = section.footer
        odd_alignments = [
            para.paragraph_format.alignment
            for para in odd_footer.paragraphs
        ]
        odd_has_right = any(a == WD_PARAGRAPH_ALIGNMENT.RIGHT for a in odd_alignments)
        if odd_has_right:
            print("PASS: Component 5 -- Odd-page footer paragraph is RIGHT-aligned (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 5 -- Odd-page footer is NOT right-aligned; found alignments: {}".format(odd_alignments))
    except Exception as e:
        print("ERROR: Component 5 -- Could not check odd-page footer alignment: {}".format(e))

    final_score = min(total_score, 1.0)
    print("\nScore: {}/1.0".format(total_score))
    print("REWARD: {}".format(final_score))
    return final_score


# Default: test against the task file path on the VM
file_path = '{}/{}'.format(WORKDIR, FILE_NAME)
if not os.path.exists(file_path):
    print("File not found: {}".format(file_path))
    print("REWARD: 0.0")
else:
    verify_task(file_path)
