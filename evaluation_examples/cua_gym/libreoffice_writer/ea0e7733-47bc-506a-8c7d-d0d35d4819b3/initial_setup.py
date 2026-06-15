"""
Initial Setup: Add page numbers in footer for book-style layout
Task ID: writer_page_068
Domain: libreoffice_writer

Creates a 16-page reference manual (A4, portrait, mirrored margins) with NO footer.
The agent must add footer with page numbers: left-aligned on even pages, right-aligned on odd pages.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_page_068'
OUTPUT = f'{WORKDIR}/Desktop/reference_manual.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a manual page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    return para


def create_initial():
    doc = Document()

    # --- Page setup: A4, portrait, mirrored margins ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    # Mirrored margins: inner=3.0cm, outer=2.0cm
    # For mirrored margins in DOCX: left = inner (odd page inside), right = outer
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # Enable mirrored margins (mirror_margins attribute in sectPr)
    sectPr = section._sectPr
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.set(qn('w:mirrorMargins'), '1')
    # Set mirrorMargins at sectPr level
    mirrorMargins = OxmlElement('w:mirrorMargins')
    existing = sectPr.find(qn('w:mirrorMargins'))
    if existing is None:
        sectPr.append(mirrorMargins)

    # NO header, NO footer in initial state
    # Ensure footer is not linked
    header = section.header
    footer = section.footer
    # Leave footer paragraphs empty (no content)
    for para in footer.paragraphs:
        for run in para.runs:
            run.clear()
        para.clear()

    # --- Page 1: Title Page ---
    doc.add_heading('Python Reference Manual', level=0)
    doc.add_paragraph()
    title_para = doc.add_paragraph('Comprehensive Developer Reference')
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_para = doc.add_paragraph('Version 3.12 | Edition 2025')
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    author_para = doc.add_paragraph('Technical Documentation Team')
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para = doc.add_paragraph('Published: January 2025')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_break(doc)

    # --- Page 2: Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('Chapter 1: Introduction to Python', '3'),
        ('Chapter 2: Data Types and Structures', '4'),
        ('Chapter 3: Control Flow', '5'),
        ('Chapter 4: Functions and Modules', '6'),
        ('Chapter 5: Object-Oriented Programming', '7'),
        ('Chapter 6: File I/O and Exceptions', '8'),
        ('Chapter 7: Standard Library Overview', '9'),
        ('Chapter 8: Concurrency and Parallelism', '10'),
        ('Chapter 9: Testing and Debugging', '11'),
        ('Chapter 10: Performance Optimization', '12'),
        ('Chapter 11: Security Best Practices', '13'),
        ('Chapter 12: Packaging and Distribution', '14'),
        ('Appendix A: Built-in Functions Reference', '15'),
        ('Appendix B: Reserved Keywords', '16'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run('\t' + page)
    add_page_break(doc)

    # --- Page 3: Chapter 1 - Introduction ---
    doc.add_heading('Chapter 1: Introduction to Python', level=1)
    doc.add_paragraph(
        'Python is a high-level, interpreted programming language with dynamic semantics. '
        'Its high-level built-in data structures, combined with dynamic typing and dynamic binding, '
        'make it very attractive for Rapid Application Development, as well as for use as a scripting '
        'or glue language to connect existing components together.'
    )
    doc.add_paragraph(
        'Python\'s simple, easy-to-learn syntax emphasizes readability and therefore reduces the cost '
        'of program maintenance. Python supports modules and packages, which encourages program modularity '
        'and code reuse. The Python interpreter and the extensive standard library are available in source '
        'or binary form without charge for all major platforms, and can be freely distributed.'
    )
    doc.add_heading('1.1 History and Design Philosophy', level=2)
    doc.add_paragraph(
        'Python was conceived in the late 1980s by Guido van Rossum at Centrum Wiskunde & Informatica (CWI) '
        'in the Netherlands as a successor to the ABC programming language, which itself was inspired by SETL. '
        'Its implementation began in December 1989. Van Rossum\'s long influence on Python is reflected in '
        'the title given to him by the Python community: Benevolent Dictator For Life (BDFL).'
    )
    doc.add_heading('1.2 Installation and Setup', level=2)
    doc.add_paragraph(
        'Python can be downloaded from the official website at python.org. The installer is available for '
        'Windows, macOS, and most Linux distributions. For development environments, it is recommended to '
        'use virtual environments to isolate project dependencies.'
    )
    add_page_break(doc)

    # --- Page 4: Chapter 2 ---
    doc.add_heading('Chapter 2: Data Types and Structures', level=1)
    doc.add_paragraph(
        'Python provides several built-in data types that cover most programming needs. Understanding '
        'these types and their performance characteristics is essential for writing efficient Python code.'
    )
    doc.add_heading('2.1 Numeric Types', level=2)
    doc.add_paragraph(
        'Python supports three numeric types: integers (int), floating-point numbers (float), and '
        'complex numbers (complex). Integers in Python 3 have arbitrary precision, meaning they can '
        'represent numbers of any size limited only by available memory.'
    )
    doc.add_heading('2.2 Sequences', level=2)
    doc.add_paragraph(
        'Sequence types include lists, tuples, and ranges. Lists are mutable ordered collections, '
        'tuples are immutable ordered collections, and ranges represent immutable sequences of numbers. '
        'Strings are also sequence types in Python, allowing for character-by-character iteration.'
    )
    doc.add_heading('2.3 Mapping Types', level=2)
    doc.add_paragraph(
        'Dictionaries are the primary mapping type in Python, providing key-value pair storage with '
        'O(1) average lookup time. As of Python 3.7, dictionaries maintain insertion order, making '
        'them predictable for iteration.'
    )
    add_page_break(doc)

    # --- Page 5: Chapter 3 ---
    doc.add_heading('Chapter 3: Control Flow', level=1)
    doc.add_paragraph(
        'Control flow statements determine the order in which instructions are executed. Python provides '
        'several control flow tools: if/elif/else conditionals, for and while loops, break, continue, '
        'pass statements, and match-case pattern matching introduced in Python 3.10.'
    )
    doc.add_heading('3.1 Conditional Statements', level=2)
    doc.add_paragraph(
        'The if statement is the most fundamental control flow tool. Python uses indentation to define '
        'code blocks, eliminating the need for braces or explicit block terminators. Conditional '
        'expressions (ternary operators) provide a concise way to express simple conditions.'
    )
    doc.add_heading('3.2 Iteration', level=2)
    doc.add_paragraph(
        'Python\'s for loop iterates over any iterable object: lists, tuples, strings, dictionaries, '
        'sets, generators, and more. The built-in range() function generates arithmetic progressions. '
        'List comprehensions, dictionary comprehensions, and set comprehensions provide concise ways '
        'to create collections based on iteration.'
    )
    add_page_break(doc)

    # --- Page 6: Chapter 4 ---
    doc.add_heading('Chapter 4: Functions and Modules', level=1)
    doc.add_paragraph(
        'Functions are reusable blocks of code that perform specific tasks. Python supports first-class '
        'functions, meaning functions can be passed as arguments, returned from other functions, and '
        'assigned to variables. This enables powerful functional programming patterns.'
    )
    doc.add_heading('4.1 Function Definition', level=2)
    doc.add_paragraph(
        'Functions are defined using the def keyword, followed by the function name and parenthesized '
        'list of formal parameters. Python supports positional arguments, keyword arguments, default '
        'parameter values, variable-length argument lists (*args and **kwargs), and keyword-only arguments.'
    )
    doc.add_heading('4.2 Lambda Functions', level=2)
    doc.add_paragraph(
        'Lambda expressions create anonymous functions using a single expression. They are syntactically '
        'restricted to a single expression and are commonly used with higher-order functions like map(), '
        'filter(), and sorted(). While convenient, they should be used sparingly to maintain readability.'
    )
    add_page_break(doc)

    # --- Page 7: Chapter 5 ---
    doc.add_heading('Chapter 5: Object-Oriented Programming', level=1)
    doc.add_paragraph(
        'Python fully supports object-oriented programming through classes and objects. OOP concepts '
        'including encapsulation, inheritance, polymorphism, and abstraction are all supported. '
        'Python also supports multiple inheritance, making it highly flexible for complex designs.'
    )
    doc.add_heading('5.1 Class Definition', level=2)
    doc.add_paragraph(
        'Classes are defined using the class keyword. The __init__ method serves as the constructor, '
        'initializing object attributes. All instance methods receive self as their first parameter, '
        'providing access to the instance\'s attributes and methods.'
    )
    doc.add_heading('5.2 Inheritance', level=2)
    doc.add_paragraph(
        'Python supports single and multiple inheritance. The super() function allows calling methods '
        'from parent classes, enabling cooperative multiple inheritance through the Method Resolution '
        'Order (MRO). Abstract base classes can be defined using the abc module.'
    )
    add_page_break(doc)

    # --- Page 8: Chapter 6 ---
    doc.add_heading('Chapter 6: File I/O and Exceptions', level=1)
    doc.add_paragraph(
        'File input/output operations in Python use file objects returned by the open() function. '
        'Python\'s exception handling mechanism provides a robust way to handle runtime errors and '
        'unexpected conditions through try/except/else/finally blocks.'
    )
    doc.add_heading('6.1 File Operations', level=2)
    doc.add_paragraph(
        'The open() function returns a file object and accepts a filename and mode parameter. Common '
        'modes include r (read), w (write), a (append), and b (binary). The with statement ensures '
        'proper resource management, automatically closing files even if an exception occurs.'
    )
    doc.add_heading('6.2 Exception Handling', level=2)
    doc.add_paragraph(
        'Python uses a hierarchy of exception classes. The BaseException class is the root, with '
        'Exception being the base for all non-system-exiting exceptions. Custom exceptions can be '
        'created by subclassing Exception, providing domain-specific error information.'
    )
    add_page_break(doc)

    # --- Page 9: Chapter 7 ---
    doc.add_heading('Chapter 7: Standard Library Overview', level=1)
    doc.add_paragraph(
        'Python\'s standard library contains over 200 modules covering string processing, data types, '
        'numeric and mathematical modules, functional programming tools, file and directory access, '
        'data compression, cryptographic services, operating system interfaces, and much more.'
    )
    doc.add_heading('7.1 String and Text Processing', level=2)
    doc.add_paragraph(
        'The re module provides regular expression operations for advanced text manipulation. The '
        'string module contains common string operations and constants. The textwrap module provides '
        'convenient functions for text formatting, such as wrapping and indentation.'
    )
    doc.add_heading('7.2 Data Persistence', level=2)
    doc.add_paragraph(
        'The pickle module serializes Python objects to byte streams. The json module encodes and '
        'decodes JSON data. The csv module implements classes for reading and writing tabular data '
        'in CSV format, compatible with spreadsheet applications.'
    )
    add_page_break(doc)

    # --- Page 10: Chapter 8 ---
    doc.add_heading('Chapter 8: Concurrency and Parallelism', level=1)
    doc.add_paragraph(
        'Python provides several approaches to concurrent and parallel execution. The threading module '
        'supports thread-based parallelism, the multiprocessing module supports process-based parallelism '
        'bypassing the GIL, and the asyncio module provides asynchronous I/O for high-performance I/O-bound tasks.'
    )
    doc.add_heading('8.1 Threading', level=2)
    doc.add_paragraph(
        'The threading module creates and manages threads in Python. Due to the Global Interpreter Lock '
        '(GIL), CPU-bound operations do not see significant performance improvements with threading. '
        'However, I/O-bound operations benefit greatly from multi-threading as threads release the GIL '
        'during I/O operations.'
    )
    doc.add_heading('8.2 Asyncio', level=2)
    doc.add_paragraph(
        'The asyncio module provides a framework for writing concurrent code using the async/await syntax. '
        'It is particularly suited for I/O-bound tasks and network applications. Event loops, coroutines, '
        'tasks, and futures are the fundamental building blocks of asyncio-based programs.'
    )
    add_page_break(doc)

    # --- Page 11: Chapter 9 ---
    doc.add_heading('Chapter 9: Testing and Debugging', level=1)
    doc.add_paragraph(
        'Testing is a critical aspect of software development. Python provides the unittest framework '
        'as part of the standard library, and the popular pytest framework offers additional features. '
        'The pdb module provides an interactive source code debugger.'
    )
    doc.add_heading('9.1 Unit Testing', level=2)
    doc.add_paragraph(
        'The unittest module provides a test framework inspired by JUnit. Test cases are created by '
        'subclassing unittest.TestCase and implementing test methods prefixed with test_. Test suites '
        'group multiple test cases, and test runners execute them and report results.'
    )
    doc.add_heading('9.2 Test Coverage', level=2)
    doc.add_paragraph(
        'Code coverage measures the proportion of source code executed during testing. The coverage.py '
        'tool integrates with pytest and unittest to generate detailed coverage reports, helping identify '
        'untested code paths and ensuring comprehensive test suites.'
    )
    add_page_break(doc)

    # --- Page 12: Chapter 10 ---
    doc.add_heading('Chapter 10: Performance Optimization', level=1)
    doc.add_paragraph(
        'Python provides various tools and techniques for measuring and improving performance. '
        'Profiling identifies bottlenecks, while optimization strategies include algorithmic improvements, '
        'caching, JIT compilation with PyPy, and using C extensions for computationally intensive tasks.'
    )
    doc.add_heading('10.1 Profiling', level=2)
    doc.add_paragraph(
        'The cProfile module provides deterministic profiling of Python programs. The pstats module '
        'analyzes profiling data. For line-by-line profiling, the line_profiler package provides '
        'detailed timing information for each line of code.'
    )
    doc.add_heading('10.2 Caching and Memoization', level=2)
    doc.add_paragraph(
        'The functools.lru_cache decorator implements memoization by caching the results of function '
        'calls. The functools.cache decorator (Python 3.9+) provides an unbounded cache. External '
        'caching solutions like Redis or Memcached are used for distributed applications.'
    )
    add_page_break(doc)

    # --- Page 13: Chapter 11 ---
    doc.add_heading('Chapter 11: Security Best Practices', level=1)
    doc.add_paragraph(
        'Security is a fundamental concern in software development. Python applications must protect '
        'against common vulnerabilities including SQL injection, cross-site scripting, insecure '
        'deserialization, and exposure of sensitive data through proper input validation and sanitization.'
    )
    doc.add_heading('11.1 Input Validation', level=2)
    doc.add_paragraph(
        'Never trust user input. Always validate and sanitize data before processing. Use parameterized '
        'queries for database operations to prevent SQL injection. The secrets module generates '
        'cryptographically strong random numbers suitable for managing tokens and passwords.'
    )
    doc.add_heading('11.2 Cryptography', level=2)
    doc.add_paragraph(
        'The hashlib module provides secure hash and message digest algorithms. For password storage, '
        'use bcrypt or argon2 rather than MD5 or SHA-1. The cryptography package provides both '
        'high-level recipes and low-level interfaces to common cryptographic algorithms.'
    )
    add_page_break(doc)

    # --- Page 14: Chapter 12 ---
    doc.add_heading('Chapter 12: Packaging and Distribution', level=1)
    doc.add_paragraph(
        'Python packaging has evolved significantly with the introduction of pyproject.toml as the '
        'standardized configuration format. The Python Package Index (PyPI) hosts thousands of '
        'third-party packages that can be installed using pip, the package installer for Python.'
    )
    doc.add_heading('12.1 Creating Packages', level=2)
    doc.add_paragraph(
        'A Python package is a directory containing an __init__.py file. The pyproject.toml file '
        'specifies build system requirements and package metadata. Build backends like setuptools, '
        'flit, and hatchling handle the packaging process.'
    )
    doc.add_heading('12.2 Virtual Environments', level=2)
    doc.add_paragraph(
        'The venv module creates lightweight virtual environments with isolated Python installations. '
        'Each environment has its own Python binary and can have its own independent set of installed '
        'packages. This prevents dependency conflicts between projects.'
    )
    add_page_break(doc)

    # --- Page 15: Appendix A ---
    doc.add_heading('Appendix A: Built-in Functions Reference', level=1)
    doc.add_paragraph(
        'Python provides numerous built-in functions available without importing any module. '
        'These functions cover type conversion, mathematical operations, I/O, sequence manipulation, '
        'object inspection, and more.'
    )
    builtin_funcs = [
        ('abs(x)', 'Return the absolute value of a number.'),
        ('all(iterable)', 'Return True if all elements of the iterable are true.'),
        ('any(iterable)', 'Return True if any element of the iterable is true.'),
        ('bin(x)', 'Convert an integer to a binary string prefixed with 0b.'),
        ('bool(x)', 'Return a Boolean value.'),
        ('callable(object)', 'Return True if the object appears callable.'),
        ('chr(i)', 'Return the string representing a character whose Unicode code point is i.'),
        ('dir([object])', 'Return the list of names in the current local scope or object attributes.'),
        ('divmod(a, b)', 'Return a pair (quotient, remainder) when dividing a by b.'),
        ('enumerate(iterable)', 'Return an enumerate object with index and value pairs.'),
        ('eval(expression)', 'Evaluate the given expression string.'),
        ('filter(function, iterable)', 'Construct an iterator from elements for which function returns true.'),
        ('format(value)', 'Convert a value to a formatted representation.'),
        ('getattr(object, name)', 'Return the value of the named attribute of object.'),
        ('hasattr(object, name)', 'Return True if the object has the named attribute.'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Function'
    hdr[1].text = 'Description'
    for func, desc in builtin_funcs:
        row = table.add_row().cells
        row[0].text = func
        row[1].text = desc
    add_page_break(doc)

    # --- Page 16: Appendix B ---
    doc.add_heading('Appendix B: Reserved Keywords', level=1)
    doc.add_paragraph(
        'The following identifiers are used as reserved words, or keywords of the language, and cannot '
        'be used as ordinary identifiers. They must be spelled exactly as written here:'
    )
    keywords = [
        'False', 'None', 'True', 'and', 'as', 'assert', 'async', 'await',
        'break', 'class', 'continue', 'def', 'del', 'elif', 'else', 'except',
        'finally', 'for', 'from', 'global', 'if', 'import', 'in', 'is',
        'lambda', 'nonlocal', 'not', 'or', 'pass', 'raise', 'return', 'try',
        'while', 'with', 'yield', 'match', 'case', 'type',
    ]
    # Display keywords in a table
    kw_table = doc.add_table(rows=1, cols=6)
    kw_table.style = 'Table Grid'
    for i, kw in enumerate(keywords):
        row_idx = i // 6
        col_idx = i % 6
        if row_idx >= len(kw_table.rows):
            kw_table.add_row()
        kw_table.rows[row_idx].cells[col_idx].text = kw

    doc.add_paragraph()
    closing = doc.add_paragraph(
        'For the most up-to-date keyword list, use the keyword module: '
        'import keyword; print(keyword.kwlist)'
    )

    # Ensure Desktop directory exists
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
