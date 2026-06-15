"""
Initial Setup: Research paper with 5 paragraphs; last paragraph is retracted conclusion
Task ID: osworld_writer_strikethrough_last_para_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Advancing Machine Learning Interpretability in Clinical Decision Support Systems', level=1)
    title.paragraph_format.space_after = Pt(12)

    # Paragraph 1: Introduction
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'Introduction. The integration of machine learning (ML) algorithms into clinical decision support systems '
        'has accelerated significantly over the past decade. As healthcare institutions increasingly rely on '
        'predictive models for diagnosis and treatment recommendations, the demand for interpretable and '
        'transparent AI systems has grown commensurately. This paper investigates the current landscape of '
        'ML interpretability techniques and their applicability within high-stakes clinical environments, '
        'drawing on evidence from twelve major hospital deployments across North America and Europe between '
        '2019 and 2024.'
    )
    intro_run.font.size = Pt(12)
    intro.paragraph_format.space_after = Pt(10)

    # Paragraph 2: Literature Review
    lit_review = doc.add_paragraph()
    lit_run = lit_review.add_run(
        'Literature Review. Prior research has established a substantial body of knowledge regarding the '
        'tension between model accuracy and interpretability. Ribeiro et al. (2016) introduced LIME, a '
        'technique enabling local explanations for any classifier, while Lundberg and Lee (2017) proposed '
        'SHAP values grounded in cooperative game theory. Subsequent work by Caruana et al. (2020) '
        'demonstrated that intelligible models such as Generalized Additive Models (GAMs) could match '
        'the performance of black-box ensembles on structured clinical data, challenging the common '
        'assumption that interpretability necessarily entails a sacrifice in predictive power. A meta-analysis '
        'by Holzinger et al. (2022) reviewed 78 studies and found that clinician trust in AI recommendations '
        'increased by 34% when model rationale was made transparent.'
    )
    lit_run.font.size = Pt(12)
    lit_review.paragraph_format.space_after = Pt(10)

    # Paragraph 3: Methodology
    method = doc.add_paragraph()
    method_run = method.add_run(
        'Methodology. This study employed a mixed-methods approach combining quantitative performance '
        'benchmarking with qualitative clinician feedback surveys. For quantitative analysis, we evaluated '
        'seven ML models—including gradient boosting (XGBoost, LightGBM), deep neural networks, and '
        'logistic regression baselines—across three clinical datasets: the MIMIC-III critical care database '
        '(n=46,520 patients), the UK Biobank cardiovascular cohort (n=22,138), and a proprietary oncology '
        'dataset from Princess Margaret Hospital (n=8,741). Model performance was assessed using AUROC, '
        'F1-score, and Brier score. Interpretability was quantified via the Fidelity-Comprehensibility '
        'Index (FCI) proposed by Martinez-Plumed et al. (2021). Qualitative data were collected via '
        'structured interviews with 47 clinicians across six specialties.'
    )
    method_run.font.size = Pt(12)
    method.paragraph_format.space_after = Pt(10)

    # Paragraph 4: Results and Discussion
    results = doc.add_paragraph()
    results_run = results.add_run(
        'Results and Discussion. Empirical results indicate that gradient boosting models achieved the '
        'highest AUROC scores (mean 0.891, SD 0.043) on the MIMIC-III dataset, while GAM-based models '
        'demonstrated comparable performance (AUROC 0.874, SD 0.038) with substantially higher FCI '
        'scores (p < 0.001). Clinician survey data revealed that 72% of participants reported greater '
        'confidence in treatment recommendations when accompanied by SHAP-based explanations, and '
        'diagnostic accuracy among participants using interpretable models improved by 11.3% compared '
        'to those using black-box outputs alone. These findings align with the growing consensus that '
        'interpretability frameworks can be integrated without materially compromising model efficacy, '
        'provided that deployment contexts are carefully considered and domain-specific calibration '
        'protocols are applied.'
    )
    results_run.font.size = Pt(12)
    results.paragraph_format.space_after = Pt(10)

    # Paragraph 5 (last): Retracted Conclusion — NO strikethrough, NO gray color in initial state
    conclusion = doc.add_paragraph()
    concl_run = conclusion.add_run(
        'Conclusion. This research conclusively demonstrates that interpretable machine learning models '
        'represent the definitive solution for all clinical decision support applications, rendering '
        'further investigation into black-box methodologies unnecessary. The universal adoption of '
        'GAM-based systems across all healthcare institutions is strongly recommended with immediate '
        'effect, as the evidence presented herein is deemed sufficient to preclude the need for '
        'additional validation studies or longitudinal follow-up. These findings have been submitted '
        'for expedited regulatory approval and should be treated as clinical guidelines pending formal '
        'review. Any contrary evidence published after the submission date of this manuscript should '
        'be disregarded in favour of the conclusions stated here.'
    )
    concl_run.font.size = Pt(12)
    # NOTE: No strikethrough and no gray color — task requires agent to add these
    conclusion.paragraph_format.space_after = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
