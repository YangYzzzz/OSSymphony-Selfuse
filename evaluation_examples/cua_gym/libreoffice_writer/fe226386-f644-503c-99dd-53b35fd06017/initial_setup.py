"""
Initial Setup: Email signature template - plain text, unformatted
Task ID: writer_mktg_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'email_signature_template'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default empty paragraph if present, then add our content
    # Clear existing paragraphs and add the single-line signature block
    # All on one line, 12pt, no special formatting
    para = doc.add_paragraph()
    run = para.add_run(
        'Rachel Kim | Senior Marketing Manager | Orion Digital | '
        'rkim@oriondigital.com | (212) 555-0147 | oriondigital.com'
    )
    run.font.size = Pt(12)
    # No bold, no color, no special formatting

    # Remove the default empty paragraph that Document() creates
    # (it's the first paragraph; our content is the second)
    default_para = doc.paragraphs[0]
    if default_para.text == '':
        p = default_para._element
        p.getparent().remove(p)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
