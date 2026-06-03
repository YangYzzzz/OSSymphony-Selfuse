"""
Reward Script: Email Signature Template Formatting
Task ID: writer_mktg_018
Domain: libreoffice_writer
Scoring:
  Component 1: Horizontal line separator (pBdr bottom border on first paragraph) — 0.20
  Component 2: 'Rachel Kim' in 12pt bold — 0.20
  Component 3: 'Senior Marketing Manager' in 11pt regular — 0.15
  Component 4: 'Orion Digital' in 11pt bold, blue (#1565C0) — 0.20
  Component 5: Contact details in 9pt, gray (#666666) — 0.15
  Component 6: Website in 9pt, blue (#1565C0) — 0.10
Total: 1.0
"""

import os
from math import sqrt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_mktg_018'
FILE_PATH = f'{WORKDIR}/email_signature_template.docx'


def color_distance(rgb1, rgb2):
    """Euclidean RGB distance."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(rgb1, rgb2)))


def has_paragraph_border(para):
    """Check if a paragraph has a bottom border via pBdr XML."""
    try:
        xml = para._element.xml
        return 'w:pBdr' in xml and ('w:bottom' in xml or 'w:top' in xml)
    except Exception:
        return False


def find_para_with_text(doc, text):
    """Find paragraph whose .text strip matches text (case-insensitive)."""
    for para in doc.paragraphs:
        if para.text.strip().lower() == text.lower():
            return para
    return None


def verify_task(file_path):
    """
    Verify email signature template formatting with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: must have at least 2 paragraphs (signature is structured)
    if len(doc.paragraphs) < 2:
        print("FAIL: Document has fewer than 2 paragraphs — signature not restructured")
        print("Score: 0.0/1.0")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Horizontal separator line — paragraph with bottom border (0.20 pts)
    # -------------------------------------------------------------------------
    try:
        paras_with_border = [p for p in doc.paragraphs if has_paragraph_border(p)]
        if len(paras_with_border) > 0:
            print("PASS: Component 1 — Horizontal separator (paragraph bottom border) found (0.20 pts)")
            total_score += 0.20
        else:
            print("FAIL: Component 1 — No paragraph with pBdr bottom border found; expected thin horizontal line separator")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: 'Rachel Kim' paragraph in 12pt bold (0.20 pts)
    # -------------------------------------------------------------------------
    try:
        para_rk = find_para_with_text(doc, 'Rachel Kim')
        if para_rk is None:
            print("FAIL: Component 2 — Paragraph with text 'Rachel Kim' not found")
        else:
            runs = [r for r in para_rk.runs if r.text.strip()]
            if not runs:
                print("FAIL: Component 2 — 'Rachel Kim' paragraph has no runs")
            else:
                run = runs[0]
                size_ok = run.font.size is not None and abs(run.font.size.pt - 12.0) < 0.5
                bold_ok = run.font.bold is True
                if size_ok and bold_ok:
                    print(f"PASS: Component 2 — 'Rachel Kim' is 12pt bold (size={run.font.size.pt}pt, bold={run.font.bold}) (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"FAIL: Component 2 — 'Rachel Kim': size_ok={size_ok} ({run.font.size.pt if run.font.size else 'None'}pt), bold_ok={bold_ok}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: 'Senior Marketing Manager' in 11pt regular (not bold) (0.15 pts)
    # -------------------------------------------------------------------------
    try:
        para_sm = find_para_with_text(doc, 'Senior Marketing Manager')
        if para_sm is None:
            print("FAIL: Component 3 — Paragraph 'Senior Marketing Manager' not found")
        else:
            runs = [r for r in para_sm.runs if r.text.strip()]
            if not runs:
                print("FAIL: Component 3 — 'Senior Marketing Manager' paragraph has no runs")
            else:
                run = runs[0]
                size_ok = run.font.size is not None and abs(run.font.size.pt - 11.0) < 0.5
                # bold should be False or None (not True)
                not_bold = run.font.bold is not True
                if size_ok and not_bold:
                    print(f"PASS: Component 3 — 'Senior Marketing Manager' is 11pt regular (size={run.font.size.pt}pt, bold={run.font.bold}) (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 3 — 'Senior Marketing Manager': size_ok={size_ok} ({run.font.size.pt if run.font.size else 'None'}pt), not_bold={not_bold} (bold={run.font.bold})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: 'Orion Digital' in 11pt bold, brand blue (#1565C0) (0.20 pts)
    # -------------------------------------------------------------------------
    try:
        para_od = find_para_with_text(doc, 'Orion Digital')
        if para_od is None:
            print("FAIL: Component 4 — Paragraph 'Orion Digital' not found")
        else:
            runs = [r for r in para_od.runs if r.text.strip()]
            if not runs:
                print("FAIL: Component 4 — 'Orion Digital' paragraph has no runs")
            else:
                run = runs[0]
                size_ok = run.font.size is not None and abs(run.font.size.pt - 11.0) < 0.5
                bold_ok = run.font.bold is True
                # Check color is near #1565C0 = (21, 101, 192)
                target_blue = (0x15, 0x65, 0xC0)
                color_ok = False
                color_actual = None
                try:
                    if run.font.color and run.font.color.type is not None:
                        rgb = run.font.color.rgb
                        color_actual = (rgb[0], rgb[1], rgb[2])
                        dist = color_distance(color_actual, target_blue)
                        color_ok = dist < 30
                except Exception:
                    pass
                if size_ok and bold_ok and color_ok:
                    print(f"PASS: Component 4 — 'Orion Digital' is 11pt bold blue (size={run.font.size.pt}pt, bold={run.font.bold}, color={color_actual}) (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"FAIL: Component 4 — 'Orion Digital': size_ok={size_ok}, bold_ok={bold_ok}, color_ok={color_ok} (color={color_actual}, target={target_blue})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Contact details line 'rkim@oriondigital.com | (212) 555-0147' in 9pt gray (#666666) (0.15 pts)
    # -------------------------------------------------------------------------
    try:
        para_contact = None
        for para in doc.paragraphs:
            if 'rkim@oriondigital.com' in para.text:
                para_contact = para
                break
        if para_contact is None:
            print("FAIL: Component 5 — Contact details paragraph not found")
        else:
            runs = [r for r in para_contact.runs if r.text.strip()]
            if not runs:
                print("FAIL: Component 5 — Contact details paragraph has no runs")
            else:
                run = runs[0]
                size_ok = run.font.size is not None and abs(run.font.size.pt - 9.0) < 0.5
                # Check color near #666666 = (102, 102, 102)
                target_gray = (0x66, 0x66, 0x66)
                color_ok = False
                color_actual = None
                try:
                    if run.font.color and run.font.color.type is not None:
                        rgb = run.font.color.rgb
                        color_actual = (rgb[0], rgb[1], rgb[2])
                        dist = color_distance(color_actual, target_gray)
                        color_ok = dist < 30
                except Exception:
                    pass
                if size_ok and color_ok:
                    print(f"PASS: Component 5 — Contact details 9pt gray (size={run.font.size.pt}pt, color={color_actual}) (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 5 — Contact details: size_ok={size_ok} ({run.font.size.pt if run.font.size else 'None'}pt), color_ok={color_ok} (color={color_actual})")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -------------------------------------------------------------------------
    # Component 6: 'oriondigital.com' in 9pt blue (#1565C0) (0.10 pts)
    # -------------------------------------------------------------------------
    try:
        para_web = find_para_with_text(doc, 'oriondigital.com')
        if para_web is None:
            print("FAIL: Component 6 — Paragraph 'oriondigital.com' not found")
        else:
            runs = [r for r in para_web.runs if r.text.strip()]
            if not runs:
                print("FAIL: Component 6 — 'oriondigital.com' paragraph has no runs")
            else:
                run = runs[0]
                size_ok = run.font.size is not None and abs(run.font.size.pt - 9.0) < 0.5
                # Check color near #1565C0 = (21, 101, 192)
                target_blue = (0x15, 0x65, 0xC0)
                color_ok = False
                color_actual = None
                try:
                    if run.font.color and run.font.color.type is not None:
                        rgb = run.font.color.rgb
                        color_actual = (rgb[0], rgb[1], rgb[2])
                        dist = color_distance(color_actual, target_blue)
                        color_ok = dist < 30
                except Exception:
                    pass
                if size_ok and color_ok:
                    print(f"PASS: Component 6 — 'oriondigital.com' 9pt blue (size={run.font.size.pt}pt, color={color_actual}) (0.10 pts)")
                    total_score += 0.10
                else:
                    print(f"FAIL: Component 6 — 'oriondigital.com': size_ok={size_ok} ({run.font.size.pt if run.font.size else 'None'}pt), color_ok={color_ok} (color={color_actual})")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Run verification
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
