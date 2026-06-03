"""
Initial Setup: Create a quarterly report with 3 comments across 3 pages.
Task ID: writer_rm_008
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from lxml import etree
import copy
from datetime import datetime

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_008'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# OOXML namespaces
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
CT_COMMENTS = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_comments_to_doc(doc, comments_data):
    """
    Add comments to a python-docx Document using direct XML manipulation.

    comments_data: list of dicts with keys:
        - id: int comment id
        - author: str
        - date: str ISO date
        - text: str comment text
        - para_index: int paragraph index to attach comment to
        - run_index: int run index within paragraph (the text to annotate)
    """
    # Step 1: Create the comments XML part
    comments_xml = etree.Element(qn('w:comments'), nsmap={
        'w': W_NS,
        'r': R_NS,
    })

    for cd in comments_data:
        comment_el = etree.SubElement(comments_xml, qn('w:comment'), {
            qn('w:id'): str(cd['id']),
            qn('w:author'): cd['author'],
            qn('w:date'): cd['date'],
        })
        # Comment content paragraph
        cp = etree.SubElement(comment_el, qn('w:p'))
        cr = etree.SubElement(cp, qn('w:r'))
        ct = etree.SubElement(cr, qn('w:t'))
        ct.text = cd['text']

    # Step 2: Add the comments part to the package
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    comments_part_name = PackURI('/word/comments.xml')
    comments_blob = etree.tostring(comments_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

    part = Part(
        partname=comments_part_name,
        content_type=CT_COMMENTS,
        blob=comments_blob,
        package=doc.part.package,
    )

    doc.part.relate_to(part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')

    # Step 3: Insert commentRangeStart, commentRangeEnd, and commentReference into document body
    for cd in comments_data:
        para = doc.paragraphs[cd['para_index']]
        para_el = para._element
        run_el = para.runs[cd['run_index']]._element

        comment_id = str(cd['id'])

        # Insert commentRangeStart before the target run
        range_start = etree.Element(qn('w:commentRangeStart'), {
            qn('w:id'): comment_id,
        })
        run_el.addprevious(range_start)

        # Insert commentRangeEnd after the target run
        range_end = etree.Element(qn('w:commentRangeEnd'), {
            qn('w:id'): comment_id,
        })
        run_el.addnext(range_end)

        # Insert a run with commentReference after commentRangeEnd
        ref_run = etree.Element(qn('w:r'))
        rpr = etree.SubElement(ref_run, qn('w:rPr'))
        rstyle = etree.SubElement(rpr, qn('w:rStyle'), {qn('w:val'): 'CommentReference'})
        comment_ref = etree.SubElement(ref_run, qn('w:commentReference'), {
            qn('w:id'): comment_id,
        })
        range_end.addnext(ref_run)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # =============================================
    # PAGE 1 - Executive Summary
    # =============================================

    # Title
    title = doc.add_heading('Q3 2025 Quarterly Business Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Analytics Corp.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.bold = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Prepared: October 15, 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # spacer

    # Executive Summary heading
    doc.add_heading('Executive Summary', level=1)

    # Para index 5 - this will get comment 0 (page 1 comment)
    p1 = doc.add_paragraph()
    run1 = p1.add_run(
        'The third quarter of 2025 demonstrated strong momentum across all business units. '
        'Total revenue reached $42.8 million, representing a 17.3% year-over-year increase. '
        'Our strategic investments in cloud infrastructure and AI-driven analytics platforms '
        'have begun yielding measurable returns, with customer acquisition costs declining by '
        '12% compared to Q2.'
    )
    run1.font.size = Pt(11)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        'Employee headcount grew to 847 full-time equivalents, with notable expansions in '
        'the Engineering and Customer Success teams. Our net promoter score improved to 72, '
        'up from 64 in the previous quarter, reflecting enhanced product quality and service '
        'delivery across all regions.'
    )
    run2.font.size = Pt(11)

    doc.add_heading('Key Highlights', level=2)

    highlights = [
        'Revenue growth of 17.3% YoY, exceeding target of 15%',
        'Operating margin improved to 23.4%, up from 20.1% in Q2',
        'Successfully launched Meridian Insights 3.0 platform',
        'Expanded into Southeast Asian markets with 3 new enterprise clients',
        'Reduced customer churn rate to 4.2% from 5.8%',
    ]
    for h in highlights:
        bp = doc.add_paragraph(h, style='List Bullet')
        for r in bp.runs:
            r.font.size = Pt(11)

    # Page break -> Page 2
    doc.add_page_break()

    # =============================================
    # PAGE 2 - Financial Performance & Charts
    # =============================================

    doc.add_heading('Financial Performance', level=1)

    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        'Revenue breakdown by segment shows continued dominance of the Enterprise Solutions '
        'division, which contributed $28.4 million (66.4% of total revenue). The SMB segment '
        'grew fastest at 24.1% YoY, reaching $8.7 million.'
    )
    run3.font.size = Pt(11)

    # Financial table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Segment', 'Q3 2025 ($M)', 'Q3 2024 ($M)', 'YoY Growth']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    data = [
        ['Enterprise Solutions', '28.4', '24.1', '17.8%'],
        ['SMB Products', '8.7', '7.0', '24.1%'],
        ['Professional Services', '3.9', '3.6', '8.3%'],
        ['Licensing & Support', '1.8', '1.8', '0.0%'],
        ['Total', '42.8', '36.5', '17.3%'],
    ]
    for r_idx, row_data in enumerate(data, 1):
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                if r_idx == 5:  # Total row
                    run.bold = True

    doc.add_paragraph()  # spacer

    # Para that gets the page 2 comment - chart reference text
    # We need to track paragraph indices carefully
    doc.add_heading('Revenue Trend Analysis', level=2)

    p_chart = doc.add_paragraph()
    run_chart = p_chart.add_run(
        'Figure 2.1: The revenue trend chart below illustrates the quarterly progression '
        'over the past eight quarters. Note the acceleration in Q2-Q3 2025 driven by '
        'enterprise deal closures and the Insights 3.0 launch.'
    )
    run_chart.font.size = Pt(11)

    p_chart2 = doc.add_paragraph()
    run_chart2 = p_chart2.add_run('[Revenue Trend Chart Placeholder - Q1 2024 through Q3 2025]')
    run_chart2.font.size = Pt(10)
    run_chart2.italic = True
    run_chart2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p_chart2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p5 = doc.add_paragraph()
    run5 = p5.add_run(
        'Cost of revenue remained stable at 34.2% of total revenue, reflecting improved '
        'operational efficiency in our cloud hosting infrastructure. Gross margin expanded '
        'to 65.8%, a 1.4 percentage point improvement over Q2 2025.'
    )
    run5.font.size = Pt(11)

    # Page break -> Page 3
    doc.add_page_break()

    # =============================================
    # PAGE 3 - Strategic Outlook & Conclusions
    # =============================================

    doc.add_heading('Strategic Outlook & Projections', level=1)

    p6 = doc.add_paragraph()
    run6 = p6.add_run(
        'Looking ahead to Q4 2025, we anticipate continued growth driven by several key '
        'initiatives. The pipeline for enterprise deals remains robust, with $18.3 million '
        'in qualified opportunities expected to close before year-end.'
    )
    run6.font.size = Pt(11)

    doc.add_heading('Q4 2025 Projections', level=2)

    projections = [
        'Revenue target: $46.2 million (+8.0% QoQ)',
        'Operating margin target: 25.0%',
        'Headcount growth: 35-40 new hires focused on Engineering and Sales',
        'Product roadmap: Meridian Insights 3.1 release with advanced NLP features',
        'Market expansion: Entry into Japanese enterprise market',
    ]
    for proj in projections:
        bp = doc.add_paragraph(proj, style='List Bullet')
        for r in bp.runs:
            r.font.size = Pt(11)

    # Conclusions paragraph that gets comment 2 (page 3)
    doc.add_heading('Conclusions', level=2)

    p_conclude = doc.add_paragraph()
    run_conclude = p_conclude.add_run(
        'Based on current trajectory and pipeline analysis, we project full-year 2025 '
        'revenue of $168.5 million, exceeding our original guidance of $155 million. '
        'The management team recommends increasing the R&D budget allocation by 15% '
        'for fiscal year 2026 to capitalize on emerging AI opportunities in the '
        'enterprise analytics space.'
    )
    run_conclude.font.size = Pt(11)

    p_final = doc.add_paragraph()
    run_final = p_final.add_run(
        'This report was prepared by the Finance and Strategy team. '
        'For questions, please contact the Office of the CFO.'
    )
    run_final.font.size = Pt(10)
    run_final.italic = True

    # Now we need to figure out paragraph indices for comment attachment
    # Let's enumerate paragraphs to find our targets
    all_paras = doc.paragraphs

    # Find paragraph indices
    exec_summary_idx = None
    chart_ref_idx = None
    conclusion_idx = None

    for i, p in enumerate(all_paras):
        text = p.text
        if 'The third quarter of 2025 demonstrated' in text:
            exec_summary_idx = i
        elif 'Figure 2.1: The revenue trend chart' in text:
            chart_ref_idx = i
        elif 'Based on current trajectory and pipeline analysis' in text:
            conclusion_idx = i

    print(f"Comment target paragraphs: exec_summary={exec_summary_idx}, chart_ref={chart_ref_idx}, conclusion={conclusion_idx}")

    # Define comments
    comments_data = [
        {
            'id': 0,
            'author': 'Sarah Lee',
            'date': '2025-10-10T09:30:00Z',
            'text': 'Great opening summary',
            'para_index': exec_summary_idx,
            'run_index': 0,
        },
        {
            'id': 1,
            'author': 'Mark Thompson',
            'date': '2025-10-11T14:15:00Z',
            'text': 'This section needs more data',
            'para_index': chart_ref_idx,
            'run_index': 0,
        },
        {
            'id': 2,
            'author': 'David Chen',
            'date': '2025-10-12T16:45:00Z',
            'text': "Let's discuss these projections",
            'para_index': conclusion_idx,
            'run_index': 0,
        },
    ]

    add_comments_to_doc(doc, comments_data)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
