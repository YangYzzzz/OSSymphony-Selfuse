"""
Initial Setup: Spelling bee practice list as a table in LibreOffice Writer
Task ID: osworld_writer_colorword_010
Domain: libreoffice_writer

Creates a .docx file with a 7x8 table (56 words), all in default black text.
Words range from 3-letter words to 13-letter words.
The file is saved as spelling_bee_list.docx per the task context.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import RGBColor, Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_010'
OUTPUT = f'{WORKDIR}/spelling_bee_list.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # 56 words for a 7x8 table - mix of short and long words
    # Short (< 8 chars): cat, beam, frog, slim, jazz, blot, quiz, trim,
    #                    stem, dusk, glow, flip, spin, clam, perk, wren,
    #                    brisk, crisp, blaze, fleet, snare, squab, prose, plank,
    #                    stomp, bland, cleft, prawn
    # Long (>= 8 chars): committee, judgment, treasure, boundary, commence,
    #                    parallel, obstacle, birthday, calendar, symphony,
    #                    champion, daughter, original, alphabet, absolute,
    #                    numerous, quantity, although, together, remember,
    #                    whenever, strength, shoulder, sentence, language,
    #                    distance, behavior, question, announce, complete
    words = [
        # Row 1 (mix)
        'cat',          'committee',    'beam',         'judgment',
        'frog',         'treasure',     'slim',         'boundary',
        # Row 2 (mix)
        'jazz',         'commence',     'blot',         'parallel',
        'quiz',         'obstacle',     'trim',         'birthday',
        # Row 3 (mix)
        'stem',         'calendar',     'dusk',         'symphony',
        'glow',         'champion',     'flip',         'daughter',
        # Row 4 (mix)
        'spin',         'original',     'clam',         'alphabet',
        'perk',         'absolute',     'wren',         'numerous',
        # Row 5 (mix)
        'brisk',        'quantity',     'crisp',        'although',
        'blaze',        'together',     'fleet',        'remember',
        # Row 6 (mix)
        'snare',        'whenever',     'squab',        'strength',
        'prose',        'shoulder',     'plank',        'sentence',
        # Row 7 (mix)
        'stomp',        'language',     'bland',        'distance',
        'cleft',        'behavior',     'prawn',        'question',
    ]

    assert len(words) == 56, f"Expected 56 words, got {len(words)}"

    doc = Document()

    # Create 7 rows x 8 columns table
    table = doc.add_table(rows=7, cols=8)
    table.style = 'Table Grid'

    idx = 0
    for row in table.rows:
        for cell in row.cells:
            word = words[idx]
            # Clear existing paragraphs content
            para = cell.paragraphs[0]
            # Clear default empty run
            para.clear()
            run = para.add_run(word)
            # All words in default black text (no color applied)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            run.font.size = Pt(12)
            idx += 1

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
