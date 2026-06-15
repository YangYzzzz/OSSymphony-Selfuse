"""
Reward Script: Color spelling bee words by length in LibreOffice Writer table
Task ID: osworld_writer_colorword_010
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): All words >= 8 characters colored dark red (#C00000)
  Component 2 (0.5): All words < 8 characters colored dark green (#006400)
"""

import os
from docx import Document
from docx.shared import RGBColor
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_010'

# Target colors (hex strings as returned by python-docx RGBColor.__str__)
DARK_RED_HEX = 'C00000'    # dark red for words >= 8 characters
DARK_GREEN_HEX = '006400'  # dark green for words < 8 characters

# Tolerance for color matching (Euclidean distance in RGB space)
# Allow slight variation in color implementation
COLOR_TOLERANCE = 30


def color_distance(rgb_str, r, g, b):
    """Compute Euclidean distance between a hex color string and an RGB tuple."""
    try:
        rr = int(rgb_str[0:2], 16)
        rg = int(rgb_str[2:4], 16)
        rb = int(rgb_str[4:6], 16)
        return sqrt((rr - r) ** 2 + (rg - g) ** 2 + (rb - b) ** 2)
    except Exception:
        return 999.0


def is_dark_red(color_str):
    """Check if color matches dark red (#C00000) within tolerance."""
    if color_str is None:
        return False
    return color_distance(color_str, 0xC0, 0x00, 0x00) <= COLOR_TOLERANCE


def is_dark_green(color_str):
    """Check if color matches dark green (#006400) within tolerance."""
    if color_str is None:
        return False
    return color_distance(color_str, 0x00, 0x64, 0x00) <= COLOR_TOLERANCE


def get_cell_color(cell):
    """Extract the font color from the first run in a cell, or None if not set."""
    for para in cell.paragraphs:
        for run in para.runs:
            try:
                if run.font.color and run.font.color.rgb:
                    return str(run.font.color.rgb)
            except Exception:
                pass
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Color words in a 7x8 spelling bee table.
      - Words >= 8 characters: dark red (#C00000)
      - Words < 8 characters: dark green (#006400)
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Verify document has exactly one table with 7 rows and 8 columns
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    total_cells = num_rows * num_cols

    print(f"INFO: Table dimensions: {num_rows} rows x {num_cols} cols = {total_cells} cells")

    # Collect all cells with their word, length, and actual color
    long_words_total = 0    # words >= 8 chars
    short_words_total = 0   # words < 8 chars
    long_words_correct = 0  # >= 8 chars correctly colored dark red
    short_words_correct = 0 # < 8 chars correctly colored dark green
    long_words_errors = []
    short_words_errors = []

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            word = cell.text.strip()
            if not word:
                continue
            color = get_cell_color(cell)
            word_len = len(word)

            if word_len >= 8:
                long_words_total += 1
                if is_dark_red(color):
                    long_words_correct += 1
                else:
                    long_words_errors.append((word, word_len, color))
            else:
                short_words_total += 1
                if is_dark_green(color):
                    short_words_correct += 1
                else:
                    short_words_errors.append((word, word_len, color))

    print(f"INFO: Long words (>=8 chars): {long_words_total} total, {long_words_correct} correctly colored dark red")
    print(f"INFO: Short words (<8 chars): {short_words_total} total, {short_words_correct} correctly colored dark green")

    # Component 1: All words >= 8 characters colored dark red (0.5 points)
    # This FAILS on initial (all black) → PASSES on golden (dark red for long words)
    try:
        if long_words_total == 0:
            print("FAIL: Component 1 — No long words (>=8 chars) found in table")
        elif long_words_correct == long_words_total:
            print(f"PASS: Component 1 — All {long_words_total} long words (>=8 chars) colored dark red #C00000 (0.5 pts)")
            total_score += 0.5
        else:
            pct = long_words_correct / long_words_total
            print(f"FAIL: Component 1 — Only {long_words_correct}/{long_words_total} long words ({pct:.0%}) correctly colored dark red")
            if long_words_errors:
                print(f"  Sample errors (first 5):")
                for w, wl, c in long_words_errors[:5]:
                    print(f"    '{w}' (len={wl}): found color={c}, expected #C00000")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All words < 8 characters colored dark green (0.5 points)
    # This FAILS on initial (all black) → PASSES on golden (dark green for short words)
    try:
        if short_words_total == 0:
            print("FAIL: Component 2 — No short words (<8 chars) found in table")
        elif short_words_correct == short_words_total:
            print(f"PASS: Component 2 — All {short_words_total} short words (<8 chars) colored dark green #006400 (0.5 pts)")
            total_score += 0.5
        else:
            pct = short_words_correct / short_words_total
            print(f"FAIL: Component 2 — Only {short_words_correct}/{short_words_total} short words ({pct:.0%}) correctly colored dark green")
            if short_words_errors:
                print(f"  Sample errors (first 5):")
                for w, wl, c in short_words_errors[:5]:
                    print(f"    '{w}' (len={wl}): found color={c}, expected #006400")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/spelling_bee_list.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
