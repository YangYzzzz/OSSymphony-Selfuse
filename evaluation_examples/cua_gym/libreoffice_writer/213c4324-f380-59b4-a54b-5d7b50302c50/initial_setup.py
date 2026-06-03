"""
Initial Setup: Argumentative essay with 5 sections for reorganization task
Task ID: wrpara_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set heading style
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0x1F, 0x37, 0x63)

    # ===== SECTION 1: Introduction (heading + 2 paragraphs) =====
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'The rapid expansion of artificial intelligence into everyday decision-making '
        'processes has sparked a vigorous debate among policymakers, technologists, and '
        'ethicists alike. As machine learning algorithms increasingly determine outcomes '
        'in healthcare diagnostics, criminal sentencing, and financial lending, the question '
        'of accountability becomes not merely academic but urgently practical. This essay '
        'argues that a comprehensive regulatory framework is essential to ensure that AI '
        'systems operate transparently and equitably across all sectors of society.'
    )

    doc.add_paragraph(
        'Proponents of unfettered AI development contend that excessive regulation stifles '
        'innovation and slows the pace of technological progress. However, the mounting '
        'evidence of algorithmic bias, opaque decision-making, and the erosion of individual '
        'privacy suggests that a laissez-faire approach is no longer tenable. By examining '
        'the historical precedents of technology regulation, analyzing current failures in '
        'AI governance, and proposing actionable policy measures, this essay seeks to chart '
        'a balanced path forward.'
    )

    # ===== SECTION 2: Background (heading + 3 paragraphs) =====
    doc.add_heading('Background', level=1)

    doc.add_paragraph(
        'The origins of artificial intelligence regulation can be traced to the early 2010s, '
        'when the European Union began drafting guidelines for data protection that would '
        'eventually culminate in the General Data Protection Regulation of 2018. The GDPR '
        'introduced the concept of a right to explanation, mandating that individuals affected '
        'by automated decisions could demand a meaningful account of the logic involved. This '
        'legislative milestone marked the first significant attempt by a governing body to '
        'impose transparency requirements on algorithmic systems.'
    )

    doc.add_paragraph(
        'In the United States, regulation has followed a more fragmented trajectory. Federal '
        'agencies such as the Federal Trade Commission and the Food and Drug Administration '
        'have issued sector-specific guidance, but no comprehensive federal AI legislation '
        'has been enacted as of 2025. Several states, most notably California and Illinois, '
        'have passed targeted laws addressing facial recognition and automated employment '
        'screening, yet these piecemeal efforts leave vast areas of AI application unregulated.'
    )

    doc.add_paragraph(
        'Meanwhile, countries in the Asia-Pacific region have adopted diverse strategies. '
        'Singapore established a voluntary Model AI Governance Framework in 2019, emphasizing '
        'industry self-regulation and best practices. China, by contrast, has pursued a more '
        'directive approach, issuing binding regulations on recommendation algorithms and '
        'deepfake technology. Japan has focused on promoting AI innovation while maintaining '
        'a light regulatory touch, relying on soft law instruments and multi-stakeholder dialogues.'
    )

    # ===== SECTION 3: Analysis (heading + 4 paragraphs) =====
    doc.add_heading('Analysis', level=1)

    doc.add_paragraph(
        'A critical examination of existing AI governance reveals three primary areas of '
        'concern: algorithmic bias, lack of transparency, and insufficient accountability '
        'mechanisms. Studies conducted by researchers at MIT and Stanford have demonstrated '
        'that commercial facial recognition systems exhibit error rates up to 34 percent '
        'higher for darker-skinned women compared to lighter-skinned men. These disparities '
        'are not merely statistical anomalies but reflect deep-seated biases embedded in '
        'training data and model architectures.'
    )

    doc.add_paragraph(
        'Transparency remains a fundamental challenge in AI governance. Many state-of-the-art '
        'models, particularly deep neural networks, function as black boxes whose internal '
        'decision-making processes resist straightforward interpretation. While techniques '
        'such as SHAP values and LIME have been developed to provide post-hoc explanations, '
        'these methods offer approximations rather than true transparency. The gap between '
        'technical explainability and meaningful public understanding poses a significant '
        'barrier to informed consent and democratic oversight.'
    )

    doc.add_paragraph(
        'The question of accountability is equally vexing. When an autonomous vehicle causes '
        'an accident or an AI-driven hiring tool discriminates against qualified applicants, '
        'the chain of responsibility is often diffuse. Developers, deployers, and end users '
        'may each bear partial responsibility, yet existing legal frameworks are ill-equipped '
        'to adjudicate liability in these novel contexts. The European Union AI Act, proposed '
        'in 2021 and finalized in 2024, attempts to address this through a risk-based '
        'classification system that assigns varying levels of obligation depending on the '
        'potential for harm.'
    )

    doc.add_paragraph(
        'Furthermore, the economic incentives driving AI deployment often conflict with '
        'regulatory objectives. Companies that invest heavily in AI systems have strong '
        'financial motivations to minimize compliance costs and resist disclosure requirements. '
        'This tension between profit and public interest underscores the need for regulatory '
        'mechanisms that are not only technically sound but also economically viable, ensuring '
        'that compliance does not disproportionately burden smaller firms while allowing '
        'large corporations to exploit regulatory loopholes.'
    )

    # ===== SECTION 4: Counter-arguments (heading + 2 paragraphs) =====
    doc.add_heading('Counter-arguments', level=1)

    doc.add_paragraph(
        'Critics of AI regulation raise several compelling objections. First, they argue that '
        'prescriptive rules risk becoming obsolete before they are fully implemented, given '
        'the pace of technological advancement. A regulation designed to govern current '
        'generative AI capabilities may prove irrelevant within two to three years as models '
        'evolve in unforeseen directions. Second, heavy-handed regulation could drive AI '
        'research and development to jurisdictions with more permissive legal environments, '
        'resulting in a regulatory race to the bottom that undermines the very protections '
        'lawmakers seek to establish.'
    )

    doc.add_paragraph(
        'Additionally, some scholars contend that the market itself provides sufficient '
        'corrective mechanisms. Companies that deploy biased or harmful AI systems face '
        'reputational damage, consumer backlash, and potential litigation, creating natural '
        'incentives for responsible development. However, this argument overlooks the '
        'significant information asymmetry between corporations and consumers, as well as '
        'the reality that many AI harms disproportionately affect marginalized communities '
        'who lack the resources to pursue legal remedies or exert meaningful market pressure.'
    )

    # ===== SECTION 5: Conclusion (heading + 2 paragraphs) =====
    doc.add_heading('Conclusion', level=1)

    doc.add_paragraph(
        'The evidence presented in this essay compels a clear verdict: voluntary guidelines '
        'and market forces alone are insufficient to govern the deployment of artificial '
        'intelligence in high-stakes domains. A robust, adaptive regulatory framework must '
        'be established at both national and international levels to ensure that AI systems '
        'operate within bounds that respect fundamental rights and promote equitable outcomes. '
        'Such a framework should incorporate mandatory impact assessments, independent auditing '
        'requirements, and meaningful avenues for redress when algorithmic decisions cause harm.'
    )

    doc.add_paragraph(
        'Looking ahead, the challenge lies in designing regulations that are flexible enough '
        'to accommodate technological evolution while remaining firm in their commitment to '
        'core principles of fairness, transparency, and accountability. International cooperation '
        'will be essential, as AI systems increasingly operate across borders and jurisdictions. '
        'The stakes are too high and the consequences too far-reaching to leave the governance '
        'of artificial intelligence to chance or to the unilateral decisions of corporate actors '
        'motivated primarily by profit.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
