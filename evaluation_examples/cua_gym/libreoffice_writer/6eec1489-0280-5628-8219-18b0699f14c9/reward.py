"""
FINAL REWARD SCRIPT - SUCCESS
Task: For the vocabulary section, make vowel-start entries red and consonant-start entries blue.
Generated: 2025-10-14 11:36:51
Status: success
Model: azure-o3
Total Steps: 1
"""

from docx import Document
from docx.shared import RGBColor
import os
import re

"""
Reward Script: Vocabulary Colour Verification
Task: In the document’s *vocabulary* section every word that starts with a vowel
must be coloured RED and every word that starts with a consonant must be coloured
BLUE.

Verification Logic
------------------
1. Load the provided .docx file with python-docx.
2. Locate the **vocabulary section**:
   • Find the first heading whose text contains the word “vocabulary” (case-insensitive).
   • The section starts with the first paragraph *after* that heading and ends
     when the next heading is encountered or at the document’s end.
3. Examine every run inside that section.
   • Extract real words using a regex (letters only).
   • Work out the expected colour for each word (vowel ⇒ red, consonant ⇒ blue).
   • Compare the run’s font colour (RGB) to the expected colour.
4. Scoring (progressive):
   • Accuracy = correctly-coloured words / total vocabulary words.
   • Base score = accuracy (linear).
   • If *all* words are coloured and every one is correct ⇒ force score = 1.0.
5. Print detailed diagnostics and final score as “REWARD: X.X”.

Security / Compliance
---------------------
• Uses only safe libraries (python-docx, re, os).
• No subprocess calls. No hard-coded success values.
• All points are awarded solely on actual verification of colours.
"""

def verify_vocab_colors(file_path: str) -> float:
    print(f"Loading document: {file_path}")

    if not os.path.exists(file_path):
        print("✗ File not found")
        print("REWARD: 0.0")
        return 0.0

    # Attempt to load the DOCX
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Error loading document: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------------------- Locate Vocabulary Section ---------------------- #
    vocab_start_idx = None
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name.lower().startswith("heading") and "vocabulary" in para.text.lower():
            vocab_start_idx = idx + 1  # section begins after this heading
            break

    # If a heading isn’t found, scan whole document (still gives a chance)
    if vocab_start_idx is None:
        print("! Vocabulary heading not found – scanning entire document")
        vocab_start_idx = 0

    # Determine where the section ends (next heading or EOF)
    vocab_end_idx = len(doc.paragraphs)
    for idx in range(vocab_start_idx, len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        style_name = para.style.name if para.style else ""
        if idx != vocab_start_idx and style_name.lower().startswith("heading"):
            vocab_end_idx = idx
            break

    # ---------------------------- Colour Check ----------------------------- #
    vowels = set("aeiou")
    rgb_red = RGBColor(255, 0, 0)
    rgb_blue = RGBColor(0, 0, 255)

    word_regex = re.compile(r"[A-Za-z]+")

    total_words = 0
    correct = 0
    incorrect = 0
    missing_colour = 0

    for para in doc.paragraphs[vocab_start_idx:vocab_end_idx]:
        for run in para.runs:
            words = word_regex.findall(run.text)
            if not words:
                continue  # Skip runs without actual words

            run_colour = run.font.color.rgb  # May be None

            for word in words:
                first_letter = word[0].lower()
                total_words += 1
                expected_colour = rgb_red if first_letter in vowels else rgb_blue

                if run_colour is None:
                    missing_colour += 1
                elif run_colour == expected_colour:
                    correct += 1
                else:
                    incorrect += 1

    # ----------------------------- Scoring --------------------------------- #
    if total_words == 0:
        print("✗ No vocabulary words found – score 0.0")
        print("REWARD: 0.0")
        return 0.0

    print(f"Total vocabulary words: {total_words}")
    print(f"✓ Correctly coloured : {correct}")
    print(f"✗ Incorrect colour   : {incorrect}")
    print(f"! Missing colour     : {missing_colour}")

    accuracy = correct / total_words
    score = accuracy  # linear mapping (progressive)

    # Grant full credit only if every word is present & colour perfect
    if incorrect == 0 and missing_colour == 0 and correct == total_words:
        score = 1.0

    # Ensure result never exceeds 1.0
    score = min(score, 1.0)

    print(f"Accuracy: {accuracy:.3f}")
    print(f"REWARD: {score}")
    return score

# ---------------------------------------------------------------------------
# Execute verification when run as a script
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    DOC_PATH = "/home/user/for_the_vocabulary_section_make_vowel_start_entries_red_and_consonant_start_entries_blue.docx"
    verify_vocab_colors(DOC_PATH)

