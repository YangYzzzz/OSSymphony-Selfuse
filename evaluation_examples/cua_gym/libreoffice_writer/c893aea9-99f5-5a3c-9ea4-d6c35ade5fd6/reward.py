"""
Reward Script: Delete the placeholder paragraph from an HR policy document.
Task ID: writer_hr_019
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Placeholder text is absent from the document
  Component 2 (0.3): Paragraph count is 34 (one fewer than original 35)
  Component 3 (0.2): Surrounding paragraphs are now adjacent (punctuality -> remote work heading)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_019'
PLACEHOLDER_TEXT = 'This section is under review and will be updated shortly.'


def persist_app_state(domain):
    """Save any unsaved LibreOffice state before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        import time
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the placeholder paragraph has been deleted from the document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_para_texts = [p.text.strip() for p in doc.paragraphs]

    # Component 1: Placeholder text is completely absent (0.5 points)
    # This is the primary verification - the exact placeholder text must not appear
    # in any paragraph of the document.
    try:
        matching_paras = [t for t in all_para_texts if PLACEHOLDER_TEXT in t]
        if len(matching_paras) == 0:
            print(f"PASS: Component 1 -- Placeholder text not found in document (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 -- Placeholder text still present in document ({len(matching_paras)} match(es))")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Paragraph count is now 34 (was 35 with the placeholder) (0.3 points)
    # The deletion of exactly one paragraph should reduce count from 35 to 34.
    try:
        para_count = len(doc.paragraphs)
        if para_count == 34:
            print(f"PASS: Component 2 -- Paragraph count is 34 as expected (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 -- Expected 34 paragraphs, found {para_count}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Surrounding context is intact -- the punctuality paragraph is
    # immediately followed by the '3.2 Remote Work Policy' heading (0.2 points)
    # In the initial file, the placeholder sat between these two paragraphs.
    try:
        adjacency_pairs = [
            (all_para_texts[i], all_para_texts[i + 1])
            for i in range(len(all_para_texts) - 1)
            if ('Repeated tardiness' in all_para_texts[i] or 'arrive at their workstations on time' in all_para_texts[i])
            and '3.2' in all_para_texts[i + 1] and 'Remote Work' in all_para_texts[i + 1]
        ]
        if len(adjacency_pairs) > 0:
            print(f"PASS: Component 3 -- Punctuality paragraph directly followed by '3.2 Remote Work Policy' heading (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 -- Surrounding paragraphs are not adjacent as expected")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
