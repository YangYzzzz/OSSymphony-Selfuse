"""
Initial Setup: HR Policy document with placeholder paragraph to be deleted.
Task ID: writer_hr_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -- Title --
    title = doc.add_heading('Greenfield Technologies Employee Policy Handbook', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Effective Date: January 1, 2025')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    version = doc.add_paragraph()
    version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version.add_run('Version 3.2 — Human Resources Department')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # -- Section 1: Introduction --
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Welcome to Greenfield Technologies. This handbook outlines the policies, '
        'procedures, and benefits that apply to all full-time and part-time employees. '
        'By accepting employment with our company, you agree to abide by the guidelines '
        'set forth in this document.'
    )
    doc.add_paragraph(
        'This handbook is not a contract of employment and does not guarantee employment '
        'for any specific duration. Greenfield Technologies reserves the right to amend, '
        'modify, or revoke any policy at any time, with or without notice.'
    )

    # -- Section 2: Code of Conduct --
    doc.add_heading('2. Code of Conduct', level=1)
    doc.add_paragraph(
        'All employees are expected to conduct themselves with professionalism, integrity, '
        'and respect for colleagues, clients, and partners. Violations of the code of conduct '
        'may result in disciplinary action, up to and including termination.'
    )

    doc.add_heading('2.1 Workplace Behavior', level=2)
    doc.add_paragraph(
        'Employees must maintain a professional demeanor at all times while on company '
        'premises or representing Greenfield Technologies at external events. Harassment, '
        'bullying, or discriminatory behavior of any kind will not be tolerated.'
    )

    doc.add_heading('2.2 Confidentiality', level=2)
    doc.add_paragraph(
        'Employees with access to proprietary information, trade secrets, or client data '
        'must safeguard such information in accordance with their confidentiality agreement. '
        'Unauthorized disclosure may result in immediate termination and legal action.'
    )

    # -- Section 3: Attendance and Work Hours --
    doc.add_heading('3. Attendance and Work Hours', level=1)
    doc.add_paragraph(
        'Standard business hours are Monday through Friday, 9:00 AM to 5:30 PM. Departments '
        'may establish alternative schedules with prior approval from the VP of Operations. '
        'Remote work arrangements are available for eligible roles as outlined in Section 3.2.'
    )

    doc.add_heading('3.1 Punctuality', level=2)
    doc.add_paragraph(
        'Employees are expected to arrive at their workstations on time. Repeated tardiness '
        'without valid justification will be documented and may lead to corrective action. '
        'Supervisors should be notified of any anticipated lateness or absence before the '
        'start of the scheduled shift.'
    )

    # *** THE PLACEHOLDER PARAGRAPH ***
    placeholder = doc.add_paragraph(
        'This section is under review and will be updated shortly.'
    )
    run = placeholder.runs[0]
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_heading('3.2 Remote Work Policy', level=2)
    doc.add_paragraph(
        'Eligible employees may request a remote work arrangement through their direct manager. '
        'Approval is contingent upon role suitability, performance history, and departmental needs. '
        'Remote employees must remain accessible during core hours (10:00 AM to 3:00 PM) and '
        'attend all mandatory in-person meetings.'
    )

    # -- Section 4: Leave Policies --
    doc.add_heading('4. Leave Policies', level=1)

    doc.add_heading('4.1 Annual Leave', level=2)
    doc.add_paragraph(
        'Full-time employees accrue 15 days of paid annual leave per calendar year. Unused leave '
        'may be carried over up to a maximum of 5 days into the following year. Employees with '
        'more than 5 years of service accrue 20 days per year.'
    )

    doc.add_heading('4.2 Sick Leave', level=2)
    doc.add_paragraph(
        'Employees are entitled to 10 days of paid sick leave per year. A medical certificate '
        'is required for absences exceeding 3 consecutive working days. Sick leave cannot be '
        'carried over and unused days do not convert to compensation.'
    )

    doc.add_heading('4.3 Parental Leave', level=2)
    doc.add_paragraph(
        'Primary caregivers are entitled to 16 weeks of paid parental leave. Secondary caregivers '
        'receive 4 weeks of paid leave. Parental leave must be taken within 12 months of the '
        "birth or adoption of a child. Employees should notify HR at least 8 weeks before the "
        'anticipated start of leave.'
    )

    # -- Section 5: Compensation and Benefits --
    doc.add_heading('5. Compensation and Benefits', level=1)
    doc.add_paragraph(
        'Greenfield Technologies is committed to providing competitive compensation packages that '
        'attract and retain top talent. Salary reviews are conducted annually in March, and '
        'adjustments are based on individual performance, market benchmarks, and company financials.'
    )

    doc.add_heading('5.1 Health Insurance', level=2)
    doc.add_paragraph(
        'All full-time employees and their dependents are eligible for comprehensive health '
        'insurance coverage, including medical, dental, and vision plans. The company subsidizes '
        '80% of premium costs. Enrollment occurs during the annual open enrollment period or '
        'within 30 days of a qualifying life event.'
    )

    doc.add_heading('5.2 Retirement Plan', level=2)
    doc.add_paragraph(
        'Employees may participate in the company 401(k) plan after 90 days of employment. '
        'Greenfield Technologies matches employee contributions up to 6% of gross salary. '
        'Vesting follows a 3-year graded schedule.'
    )

    # -- Section 6: Disciplinary Procedures --
    doc.add_heading('6. Disciplinary Procedures', level=1)
    doc.add_paragraph(
        'Greenfield Technologies follows a progressive discipline approach. Depending on the '
        'severity of the infraction, the process may include verbal warning, written warning, '
        'suspension, and termination. Serious violations such as theft, violence, or substance '
        'abuse may result in immediate termination without prior warning.'
    )

    # -- Closing --
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(24)
    run = closing.add_run(
        'This handbook was last revised on December 15, 2024. For questions or clarifications, '
        'please contact the Human Resources Department at hr@greenfieldtech.com.'
    )
    run.font.size = Pt(10)
    run.font.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
