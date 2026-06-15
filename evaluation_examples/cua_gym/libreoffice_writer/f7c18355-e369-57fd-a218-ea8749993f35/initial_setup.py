"""
Initial Setup: Merge review changes from two reviewers into base document
Task ID: writer_rm_046
Domain: libreoffice_writer

Creates three documents:
  - Report_Base.docx: 20-page original report with no tracked changes
  - Report_ReviewA.docx: Copy with 8 tracked changes from Reviewer A (grammar fixes)
  - Report_ReviewB.docx: Copy with 6 tracked changes from Reviewer B (content additions)
"""

import os
import shlex
import subprocess
import time
import copy
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_046'
BASE_OUTPUT = f'{WORKDIR}/Report_Base.docx'
REVIEW_A_OUTPUT = f'{WORKDIR}/Report_ReviewA.docx'
REVIEW_B_OUTPUT = f'{WORKDIR}/Report_ReviewB.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_base_document():
    """Create a realistic 20-page business report."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Quarterly Business Performance Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Fiscal Quarter Q3 2025 — July to September')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Financial Performance Overview',
        '3. Revenue Analysis by Segment',
        '4. Operating Expenses Breakdown',
        '5. Product Development Updates',
        '6. Market Expansion Strategy',
        '7. Human Resources and Talent Acquisition',
        '8. Customer Satisfaction Metrics',
        '9. Risk Assessment and Mitigation',
        '10. Forward-Looking Projections',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'Meridian Technologies Inc. delivered strong results in Q3 2025, with consolidated '
        'revenue reaching $247.8 million, representing a 12.3% year-over-year increase. '
        'The company achieved an operating margin of 18.7%, up from 16.2% in the prior '
        'year quarter, driven by improved operational efficiency and favorable product mix.'
    )
    doc.add_paragraph(
        'Key highlights for the quarter include the successful launch of the CloudMatrix '
        'platform, which onboarded 340 enterprise clients within its first 90 days. The '
        'Asia-Pacific region emerged as the fastest-growing market, with revenue growth of '
        '28.4% compared to the same period last year. Additionally, the company completed '
        'the acquisition of DataSync Solutions, adding complementary data integration '
        'capabilities to our product portfolio.'
    )
    doc.add_paragraph(
        'Our workforce expanded to 4,832 employees globally, with strategic hires in '
        'engineering and product management. Employee retention remained strong at 92.1%, '
        'reflecting the effectiveness of our talent engagement programs. The board approved '
        'a share repurchase program of up to $50 million, underscoring confidence in the '
        'companys long-term growth trajectory.'
    )

    doc.add_page_break()

    # Section 2: Financial Performance Overview
    doc.add_heading('2. Financial Performance Overview', level=1)
    doc.add_paragraph(
        'The financial performance of Meridian Technologies in Q3 2025 demonstrated '
        'robust growth across all major metrics. Total revenue increased to $247.8 million '
        'from $220.7 million in Q3 2024, exceeding analyst consensus estimates by '
        'approximately $8.3 million.'
    )

    doc.add_heading('2.1 Revenue Composition', level=2)
    doc.add_paragraph(
        'Software licensing revenue contributed $142.6 million, representing 57.5% of '
        'total revenue. Subscription and SaaS revenue grew to $78.9 million, a 23.1% '
        'increase reflecting the ongoing transition to recurring revenue models. '
        'Professional services accounted for the remaining $26.3 million.'
    )

    # Financial table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Metric', 'Q3 2025', 'Q3 2024', 'YoY Change']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Total Revenue', '$247.8M', '$220.7M', '+12.3%'],
        ['Software Licensing', '$142.6M', '$131.2M', '+8.7%'],
        ['SaaS/Subscription', '$78.9M', '$64.1M', '+23.1%'],
        ['Professional Services', '$26.3M', '$25.4M', '+3.5%'],
        ['Gross Margin', '72.4%', '70.1%', '+2.3pp'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_paragraph('')  # spacer

    doc.add_heading('2.2 Profitability Analysis', level=2)
    doc.add_paragraph(
        'Gross profit for the quarter was $179.4 million, yielding a gross margin of '
        '72.4%, an improvement of 230 basis points compared to the prior year. This '
        'improvement was primarily driven by higher software mix, operational leverage '
        'in cloud infrastructure, and reduced third-party licensing costs.'
    )
    doc.add_paragraph(
        'Operating income reached $46.3 million, with an operating margin of 18.7%. '
        'Net income attributable to shareholders was $34.8 million, or $1.42 per '
        'diluted share, compared to $28.1 million, or $1.15 per diluted share, in '
        'Q3 2024.'
    )

    doc.add_page_break()

    # Section 3: Revenue Analysis by Segment
    doc.add_heading('3. Revenue Analysis by Segment', level=1)
    doc.add_paragraph(
        'Meridian Technologies operates through three primary business segments: '
        'Enterprise Solutions, Cloud Infrastructure, and Digital Commerce. Each segment '
        'contributed meaningfully to the overall growth trajectory during Q3 2025.'
    )

    doc.add_heading('3.1 Enterprise Solutions', level=2)
    doc.add_paragraph(
        'The Enterprise Solutions segment generated $118.4 million in revenue, accounting '
        'for 47.8% of total revenue. This represents a 9.2% increase from the $108.4 million '
        'reported in Q3 2024. The segment benefited from strong demand for our ERP integration '
        'suite and the newly launched compliance automation module.'
    )
    doc.add_paragraph(
        'Notable enterprise wins during the quarter included contracts with Pacific Northwest '
        'Healthcare Group ($4.2 million annual contract value), Stellaris Financial Partners '
        '($3.8 million), and Meridian Aerospace Industries ($2.9 million). The average deal '
        'size increased 15% to $1.8 million, reflecting successful upmarket positioning.'
    )

    doc.add_heading('3.2 Cloud Infrastructure', level=2)
    doc.add_paragraph(
        'Cloud Infrastructure revenue surged to $89.7 million, a 19.8% increase from '
        '$74.9 million in the comparable quarter. The CloudMatrix platform launch was the '
        'primary catalyst, contributing $12.3 million in its inaugural quarter. Existing '
        'cloud customers expanded their usage by an average of 34%, driven by increased '
        'data processing and storage requirements.'
    )

    doc.add_heading('3.3 Digital Commerce', level=2)
    doc.add_paragraph(
        'Digital Commerce contributed $39.7 million, representing a 6.1% increase from '
        '$37.4 million. While growth moderated compared to previous quarters, the segment '
        'maintained healthy margins of 68.2%. The launch of our AI-powered recommendation '
        'engine drove a 22% improvement in conversion rates for participating merchants.'
    )

    doc.add_page_break()

    # Section 4: Operating Expenses Breakdown
    doc.add_heading('4. Operating Expenses Breakdown', level=1)
    doc.add_paragraph(
        'Total operating expenses for Q3 2025 were $133.1 million, representing 53.7% '
        'of revenue. While absolute expenses increased by $11.8 million year-over-year, '
        'the expense ratio improved by 160 basis points, demonstrating effective cost '
        'management alongside revenue growth.'
    )

    doc.add_heading('4.1 Research and Development', level=2)
    doc.add_paragraph(
        'R&D expenditures totaled $58.7 million, or 23.7% of revenue, compared to '
        '$52.4 million (23.8%) in Q3 2024. Key investment areas included the CloudMatrix '
        'platform enhancements, next-generation AI capabilities for the Enterprise Solutions '
        'suite, and security infrastructure improvements. The engineering headcount grew to '
        '1,847 across six development centers globally.'
    )

    doc.add_heading('4.2 Sales and Marketing', level=2)
    doc.add_paragraph(
        'Sales and marketing expenses were $49.8 million, representing 20.1% of revenue, '
        'compared to $45.2 million (20.5%) in the prior year. The slight improvement in '
        'efficiency ratio reflected higher average deal sizes and improved lead conversion '
        'rates through our digital marketing initiatives.'
    )

    doc.add_heading('4.3 General and Administrative', level=2)
    doc.add_paragraph(
        'G&A costs were $24.6 million, or 9.9% of revenue, compared to $23.7 million '
        '(10.7%) in Q3 2024. The decrease as a percentage of revenue was driven by '
        'operational efficiencies in back-office processes and the consolidation of '
        'regional administrative functions.'
    )

    doc.add_page_break()

    # Section 5: Product Development Updates
    doc.add_heading('5. Product Development Updates', level=1)
    doc.add_paragraph(
        'The product development organization achieved several significant milestones '
        'during Q3 2025, advancing both new product initiatives and enhancements to '
        'existing offerings.'
    )

    doc.add_heading('5.1 CloudMatrix Platform', level=2)
    doc.add_paragraph(
        'The CloudMatrix platform, launched in July 2025, exceeded initial adoption '
        'projections. The platform provides unified cloud orchestration across multi-cloud '
        'environments, supporting AWS, Azure, and Google Cloud Platform. Key features '
        'include automated workload balancing, predictive cost optimization, and integrated '
        'security compliance monitoring.'
    )
    doc.add_paragraph(
        'Customer feedback has been overwhelmingly positive, with a Net Promoter Score '
        'of 72 among early adopters. The product roadmap for Q4 includes Kubernetes-native '
        'deployment support, enhanced observability dashboards, and integration with major '
        'CI/CD pipelines.'
    )

    doc.add_heading('5.2 Enterprise Suite v8.0', level=2)
    doc.add_paragraph(
        'Development of Enterprise Suite v8.0 progressed on schedule, with the beta '
        'release planned for November 2025. Major new features include AI-driven workflow '
        'automation, advanced document processing capabilities, and a redesigned user '
        'interface based on extensive usability research involving over 200 enterprise users.'
    )

    doc.add_page_break()

    # Section 6: Market Expansion Strategy
    doc.add_heading('6. Market Expansion Strategy', level=1)
    doc.add_paragraph(
        'Meridian Technologies continued to execute its global expansion strategy during '
        'Q3 2025, with particular focus on the Asia-Pacific and European markets.'
    )

    doc.add_heading('6.1 Asia-Pacific Growth', level=2)
    doc.add_paragraph(
        'The Asia-Pacific region delivered exceptional results, with revenue of $52.3 million, '
        'representing 28.4% year-over-year growth. The company established new offices in '
        'Singapore and Sydney, bringing the total APAC headcount to 486. Strategic partnerships '
        'with regional system integrators, including Tata Consultancy Services and Infosys, '
        'expanded our go-to-market reach in India and Southeast Asia.'
    )

    doc.add_heading('6.2 European Market', level=2)
    doc.add_paragraph(
        'European operations generated $68.9 million in revenue, a 14.7% increase from the '
        'prior year. The GDPR-compliant product variants continued to resonate with enterprise '
        'customers, particularly in the financial services and healthcare sectors. The Frankfurt '
        'data center expansion, completed in August, added 40% more capacity to support growing '
        'demand for localized data processing.'
    )

    doc.add_heading('6.3 North American Market', level=2)
    doc.add_paragraph(
        'North America remained the largest market with revenue of $126.6 million, representing '
        '51.1% of total revenue. Growth of 7.8% reflected the mature nature of this market, '
        'though expansion into mid-market segments through simplified product offerings showed '
        'promising early results with 45 new accounts added during the quarter.'
    )

    doc.add_page_break()

    # Section 7: Human Resources
    doc.add_heading('7. Human Resources and Talent Acquisition', level=1)
    doc.add_paragraph(
        'The human capital strategy continued to focus on attracting and retaining world-class '
        'talent to support the companys growth objectives. Total headcount reached 4,832 at '
        'quarter-end, a net increase of 312 positions from Q2 2025.'
    )
    doc.add_paragraph(
        'Key HR metrics for Q3 2025 included an employee retention rate of 92.1%, compared '
        'to the industry average of 85.3%. The company processed 2,847 job applications, '
        'extending offers to 418 candidates with an acceptance rate of 84.2%. The average '
        'time-to-hire decreased to 28 days from 34 days in the prior quarter, reflecting '
        'improvements in the recruitment process.'
    )
    doc.add_paragraph(
        'The Diversity, Equity, and Inclusion initiative made progress, with women '
        'representing 38.4% of the global workforce, up from 36.1% a year ago. '
        'The company launched a new mentorship program pairing senior leaders with '
        'high-potential employees from underrepresented backgrounds, with 156 participants '
        'enrolled in the inaugural cohort.'
    )

    doc.add_page_break()

    # Section 8: Customer Satisfaction
    doc.add_heading('8. Customer Satisfaction Metrics', level=1)
    doc.add_paragraph(
        'Customer satisfaction remained a core priority, with multiple indicators showing '
        'positive trends during Q3 2025.'
    )
    doc.add_paragraph(
        'The overall Customer Satisfaction Score (CSAT) improved to 4.6 out of 5.0, up from '
        '4.4 in Q3 2024. The Net Promoter Score (NPS) reached 58, placing Meridian '
        'Technologies in the top quartile of enterprise software providers. Customer churn '
        'rate declined to 3.2% annually, the lowest in the companys history.'
    )
    doc.add_paragraph(
        'The customer support organization resolved 94.7% of tickets within the first '
        'response, and average resolution time decreased to 4.2 hours from 5.8 hours in '
        'the previous quarter. The launch of the AI-powered support assistant in August '
        'contributed to a 30% reduction in routine ticket volume.'
    )

    # Customer metrics table
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    for i, h in enumerate(['Metric', 'Q3 2025', 'Q3 2024']):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    cust_data = [
        ['CSAT Score', '4.6 / 5.0', '4.4 / 5.0'],
        ['Net Promoter Score', '58', '51'],
        ['Annual Churn Rate', '3.2%', '4.1%'],
        ['First Response Resolution', '94.7%', '91.2%'],
        ['Avg Resolution Time', '4.2 hours', '5.8 hours'],
    ]
    for r, row_data in enumerate(cust_data, 1):
        for c, val in enumerate(row_data):
            table2.rows[r].cells[c].text = val

    doc.add_page_break()

    # Section 9: Risk Assessment
    doc.add_heading('9. Risk Assessment and Mitigation', level=1)
    doc.add_paragraph(
        'The enterprise risk management framework identified several key risk factors '
        'requiring ongoing monitoring and mitigation efforts during Q3 2025.'
    )

    doc.add_heading('9.1 Cybersecurity Risks', level=2)
    doc.add_paragraph(
        'The threat landscape continued to evolve, with a 15% increase in attempted '
        'intrusions compared to the prior quarter. The security operations center '
        'successfully mitigated all attempted breaches, with zero confirmed data '
        'exfiltration events. Investment in advanced threat detection capabilities '
        'increased by $3.2 million during the quarter.'
    )

    doc.add_heading('9.2 Regulatory Compliance', level=2)
    doc.add_paragraph(
        'The regulatory environment remained complex, with new data protection '
        'regulations enacted in three additional jurisdictions. The compliance team '
        'completed gap assessments for all new requirements and implemented necessary '
        'changes ahead of enforcement deadlines. The company maintained ISO 27001, '
        'SOC 2 Type II, and GDPR compliance certifications without any findings.'
    )

    doc.add_heading('9.3 Supply Chain and Vendor Risks', level=2)
    doc.add_paragraph(
        'Third-party vendor risk assessments were conducted for all 142 critical vendors. '
        'Two vendors were identified as high-risk due to financial instability, and '
        'contingency plans were activated including identification of alternative suppliers. '
        'The company maintained a 99.97% uptime across all production services.'
    )

    doc.add_page_break()

    # Section 10: Forward-Looking Projections
    doc.add_heading('10. Forward-Looking Projections', level=1)
    doc.add_paragraph(
        'Based on current business momentum and market conditions, the company is raising '
        'its full-year 2025 guidance. Management now expects total revenue in the range of '
        '$980 million to $1.01 billion, up from the previous guidance of $950 million to '
        '$975 million.'
    )
    doc.add_paragraph(
        'For Q4 2025, the company anticipates revenue of $260 million to $270 million, '
        'reflecting continued strong demand for cloud and enterprise solutions. Operating '
        'margin is expected to remain in the range of 18% to 19%, with incremental '
        'investments in sales capacity and product development partially offsetting '
        'margin expansion from revenue growth.'
    )
    doc.add_paragraph(
        'Strategic priorities for the remainder of fiscal 2025 include the general '
        'availability launch of Enterprise Suite v8.0, expansion of the CloudMatrix '
        'partner ecosystem, and further penetration of the mid-market segment through '
        'simplified deployment options and competitive pricing.'
    )

    # Guidance table
    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Table Grid'
    for i, h in enumerate(['Metric', 'Updated Guidance', 'Previous Guidance']):
        cell = table3.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    guidance_data = [
        ['FY 2025 Revenue', '$980M - $1.01B', '$950M - $975M'],
        ['Q4 2025 Revenue', '$260M - $270M', '$245M - $255M'],
        ['Operating Margin', '18% - 19%', '17% - 18%'],
    ]
    for r, row_data in enumerate(guidance_data, 1):
        for c, val in enumerate(row_data):
            table3.rows[r].cells[c].text = val

    doc.add_paragraph('')  # spacer

    doc.add_heading('Disclaimer', level=2)
    doc.add_paragraph(
        'This report contains forward-looking statements that involve risks and uncertainties. '
        'Actual results may differ materially from those projected. Factors that could cause '
        'actual results to differ include economic conditions, competitive dynamics, technology '
        'changes, and regulatory developments. The company undertakes no obligation to update '
        'forward-looking statements.'
    )

    doc.add_paragraph('')
    doc.add_paragraph('')

    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = closing.add_run('Prepared by the Office of the Chief Financial Officer')
    run.italic = True
    run.font.size = Pt(10)

    approval = doc.add_paragraph()
    approval.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = approval.add_run('Approved for Distribution: September 30, 2025')
    run.font.size = Pt(10)

    doc.save(BASE_OUTPUT)
    print(f'Base document created: {BASE_OUTPUT}')
    return doc


def add_tracked_insertion(paragraph, text, author, date_str, run_props=None):
    """Add a tracked insertion (w:ins) to a paragraph at the end."""
    ins_elem = parse_xml(
        f'<w:ins {nsdecls("w")} '
        f'w:id="{id(text) % 10000}" '
        f'w:author="{author}" '
        f'w:date="{date_str}"/>'
    )
    r_elem = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{text}</w:t></w:r>')
    if run_props is not None:
        r_elem.insert(0, run_props)
    ins_elem.append(r_elem)
    paragraph._element.append(ins_elem)


def add_tracked_deletion(paragraph, run_index, author, date_str):
    """Mark an existing run in a paragraph as a tracked deletion (w:del)."""
    if run_index >= len(paragraph.runs):
        return
    run_elem = paragraph.runs[run_index]._element
    parent = run_elem.getparent()

    del_elem = parse_xml(
        f'<w:del {nsdecls("w")} '
        f'w:id="{(id(run_elem) + run_index) % 10000}" '
        f'w:author="{author}" '
        f'w:date="{date_str}"/>'
    )
    # Need to use delText instead of t
    del_run = copy.deepcopy(run_elem)
    for t_elem in del_run.findall(qn('w:t')):
        del_text = parse_xml(
            f'<w:delText {nsdecls("w")} xml:space="preserve">{t_elem.text or ""}</w:delText>'
        )
        t_elem.getparent().replace(t_elem, del_text)
    del_elem.append(del_run)

    parent.replace(run_elem, del_elem)


def create_review_a():
    """Create Report_ReviewA.docx with 8 tracked changes from Reviewer A (grammar fixes)."""
    import shutil
    shutil.copy(BASE_OUTPUT, REVIEW_A_OUTPUT)

    doc = Document(REVIEW_A_OUTPUT)
    author = 'Reviewer A'
    date_str = '2025-10-02T09:30:00Z'

    # We'll apply 8 grammar-fix tracked changes by manipulating the XML directly.
    # Strategy: For each change, find a paragraph, split it to isolate the word to fix,
    # then add deletion + insertion tracked changes.

    # Access the document body XML
    body = doc.element.body

    # Helper to create a del+ins pair (substitution as tracked change)
    def make_substitution(para_elem, old_text, new_text, change_id):
        """Find old_text in a run within para_elem, replace with del+ins tracked changes."""
        for r in para_elem.findall(qn('w:r')):
            t = r.find(qn('w:t'))
            if t is not None and t.text and old_text in t.text:
                full_text = t.text
                before, _, after = full_text.partition(old_text)

                parent = r.getparent()
                idx = list(parent).index(r)

                # Run with text before the change
                if before:
                    r_before = copy.deepcopy(r)
                    r_before.find(qn('w:t')).text = before
                    r_before.find(qn('w:t')).set(qn('xml:space'), 'preserve')
                    parent.insert(idx, r_before)
                    idx += 1

                # Deletion of old text
                del_elem = etree.SubElement(parent, qn('w:del'))
                del_elem.set(qn('w:id'), str(change_id))
                del_elem.set(qn('w:author'), author)
                del_elem.set(qn('w:date'), date_str)
                del_run = copy.deepcopy(r)
                del_t = del_run.find(qn('w:t'))
                del_text_elem = etree.SubElement(del_run, qn('w:delText'))
                del_text_elem.text = old_text
                del_text_elem.set(qn('xml:space'), 'preserve')
                del_run.remove(del_t)
                del_elem.append(del_run)
                parent.insert(idx, del_elem)
                idx += 1

                # Insertion of new text
                ins_elem = etree.SubElement(parent, qn('w:ins'))
                ins_elem.set(qn('w:id'), str(change_id + 1))
                ins_elem.set(qn('w:author'), author)
                ins_elem.set(qn('w:date'), date_str)
                ins_run = copy.deepcopy(r)
                ins_run.find(qn('w:t')).text = new_text
                ins_run.find(qn('w:t')).set(qn('xml:space'), 'preserve')
                ins_elem.append(ins_run)
                parent.insert(idx, ins_elem)
                idx += 1

                # Run with text after the change
                if after:
                    r_after = copy.deepcopy(r)
                    r_after.find(qn('w:t')).text = after
                    r_after.find(qn('w:t')).set(qn('xml:space'), 'preserve')
                    parent.insert(idx, r_after)
                    idx += 1

                # Remove original run
                parent.remove(r)
                return True
        return False

    # 8 grammar fixes by Reviewer A (find specific text in paragraphs)
    grammar_fixes = [
        ("companys long-term", "company's long-term", 100),
        ("companys growth objectives", "company's growth objectives", 102),
        ("companys history", "company's history", 104),
        ("newly launched compliance automation", "newly-launched compliance automation", 106),  # hyphenate
        ("next-generation AI capabilities for the Enterprise Solutions", "next-generation AI capabilities for the Enterprise Solutions'", 108),  # possessive
        ("identification of alternative suppliers", "the identification of alternative suppliers", 110),  # article
        ("contributed meaningfully to the overall", "contributed meaningfully to overall", 112),  # remove unnecessary article
        ("the expense ratio improved by 160 basis points, demonstrating", "the expense ratio improved by 160 basis points demonstrating", 114),  # remove comma
    ]

    paras = body.findall(qn('w:p'))
    applied = 0
    for old_text, new_text, cid in grammar_fixes:
        for p in paras:
            if make_substitution(p, old_text, new_text, cid):
                applied += 1
                break

    print(f'Review A: Applied {applied} tracked changes')
    doc.save(REVIEW_A_OUTPUT)
    print(f'Review A document created: {REVIEW_A_OUTPUT}')


def create_review_b():
    """Create Report_ReviewB.docx with 6 tracked changes from Reviewer B (content additions)."""
    import shutil
    shutil.copy(BASE_OUTPUT, REVIEW_B_OUTPUT)

    doc = Document(REVIEW_B_OUTPUT)
    author = 'Reviewer B'
    date_str = '2025-10-03T14:15:00Z'

    body = doc.element.body
    paras = body.findall(qn('w:p'))

    # 6 content additions by Reviewer B — insert additional sentences/clauses
    # We'll add tracked insertions at the end of specific paragraphs

    additions = [
        # (search text to find paragraph, text to insert, change_id)
        ("driven by improved operational efficiency and favorable product mix",
         " The management team attributes this success to disciplined cost control and strategic investment allocation.",
         200),
        ("onboarded 340 enterprise clients within its first 90 days",
         " This milestone exceeds the internal target of 250 clients by a significant margin.",
         202),
        ("Asia-Pacific region emerged as the fastest-growing market",
         " Management plans to allocate an additional $15 million to APAC expansion in Q4.",
         204),
        ("average deal size increased 15% to $1.8 million",
         " The sales pipeline for Q4 includes several opportunities exceeding $5 million in contract value.",
         206),
        ("Kubernetes-native deployment support",
         " and an AI-powered anomaly detection feature that leverages proprietary machine learning models",
         208),
        ("99.97% uptime across all production services",
         " This performance ranks Meridian Technologies among the top three providers in the industry for service reliability.",
         210),
    ]

    applied = 0
    for search_text, insert_text, cid in additions:
        for p in paras:
            # Check if this paragraph contains the search text
            full_text = ''.join(
                (t.text or '') for t in p.iter(qn('w:t'))
            )
            if search_text in full_text:
                # Add tracked insertion at the end of the paragraph
                ins_elem = etree.SubElement(p, qn('w:ins'))
                ins_elem.set(qn('w:id'), str(cid))
                ins_elem.set(qn('w:author'), author)
                ins_elem.set(qn('w:date'), date_str)

                ins_run = etree.SubElement(ins_elem, qn('w:r'))
                # Copy run properties from last run if available
                runs = p.findall(qn('w:r'))
                if runs:
                    rpr = runs[-1].find(qn('w:rPr'))
                    if rpr is not None:
                        ins_run.append(copy.deepcopy(rpr))

                ins_t = etree.SubElement(ins_run, qn('w:t'))
                ins_t.text = insert_text
                ins_t.set(qn('xml:space'), 'preserve')

                applied += 1
                break

    print(f'Review B: Applied {applied} tracked changes')
    doc.save(REVIEW_B_OUTPUT)
    print(f'Review B document created: {REVIEW_B_OUTPUT}')


def create_initial():
    create_base_document()
    create_review_a()
    create_review_b()

    # Launch LibreOffice Writer with the base document
    launch_gui(f'libreoffice --writer "{BASE_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with Report_Base.docx')


create_initial()
