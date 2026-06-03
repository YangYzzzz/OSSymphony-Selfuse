"""
Initial Setup: Lab Report with 4 images, no captions
Task ID: writer_pd_007
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import struct
import zlib
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_placeholder_image(path, width=640, height=400, color=(200, 200, 220), label=""):
    """Create a simple solid-color PNG image using pure Python (no Pillow encoder needed)."""
    def make_chunk(chunk_type, data):
        chunk = chunk_type + data
        return struct.pack('>I', len(data)) + chunk + struct.pack('>I', zlib.crc32(chunk) & 0xFFFFFFFF)

    # PNG signature
    signature = b'\x89PNG\r\n\x1a\n'
    # IHDR
    ihdr_data = struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0)  # 8-bit RGB
    ihdr = make_chunk(b'IHDR', ihdr_data)
    # IDAT - raw pixel data with filter byte 0 per row
    raw_data = b''
    r, g, b = color
    row_bytes = b'\x00' + bytes([r, g, b]) * width  # filter byte + RGB pixels
    for y in range(height):
        # Add a border effect: darker pixels at edges
        if y < 2 or y >= height - 2:
            border_row = b'\x00' + bytes([100, 100, 120]) * width
            raw_data += border_row
        else:
            # Left and right border
            row = bytearray(b'\x00')
            for x in range(width):
                if x < 2 or x >= width - 2:
                    row.extend([100, 100, 120])
                else:
                    row.extend([r, g, b])
            raw_data += bytes(row)
    compressed = zlib.compress(raw_data)
    idat = make_chunk(b'IDAT', compressed)
    # IEND
    iend = make_chunk(b'IEND', b'')
    with open(path, 'wb') as f:
        f.write(signature + ihdr + idat + iend)


def add_page_break(doc):
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def add_body_paragraphs(doc, text_lines, heading=None, heading_level=2):
    """Add a heading followed by body paragraphs."""
    if heading:
        doc.add_heading(heading, level=heading_level)
    for line in text_lines:
        para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Create placeholder images
    img_dir = '/tmp/lab_report_images'
    os.makedirs(img_dir, exist_ok=True)

    image_configs = [
        ("setup_photo.png", (180, 200, 220), "Experimental Setup Photo"),
        ("data_distribution.png", (200, 220, 180), "Raw Data Distribution"),
        ("analysis_results.png", (220, 200, 180), "Analysis Results Graph"),
        ("comparison_chart.png", (180, 220, 200), "Comparison Chart"),
    ]
    image_paths = []
    for fname, color, label in image_configs:
        path = os.path.join(img_dir, fname)
        create_placeholder_image(path, 640, 400, color, label)
        image_paths.append(path)

    # ===== PAGE 1: Title Page =====
    doc.add_paragraph()  # spacer
    doc.add_paragraph()
    title = doc.add_heading('Advanced Materials Characterization:\nThermal Analysis of Polymer Composites', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    meta_info = [
        "Laboratory Report - MATL 4520",
        "Department of Materials Science and Engineering",
        "Northwestern Polytechnic University",
        "",
        "Principal Investigator: Dr. Elena Vasquez",
        "Research Assistants: Kenji Tanaka, Priya Sharma, Lucas Andersson",
        "",
        "Date: March 15, 2025",
        "Report Version: 2.1 (Final)",
    ]
    for line in meta_info:
        p = doc.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if line:
            for run in p.runs:
                run.font.size = Pt(12)

    # ===== PAGE 2: Abstract & Introduction =====
    add_page_break(doc)
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive thermal analysis of carbon-fiber reinforced polymer '
        'composites (CFRPs) subjected to varying temperature gradients. Using differential scanning '
        'calorimetry (DSC) and thermogravimetric analysis (TGA), we characterized the thermal stability '
        'and phase transition behavior of three composite formulations. Our findings indicate that the '
        'addition of 2.5% nano-silica particles significantly improves the glass transition temperature '
        '(Tg) by 18.3K while maintaining mechanical integrity under cyclic thermal loading.'
    )

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Carbon-fiber reinforced polymer composites have become essential materials in aerospace, '
        'automotive, and renewable energy applications due to their exceptional strength-to-weight '
        'ratio and corrosion resistance. However, their thermal performance remains a critical '
        'limitation, particularly in environments exceeding 150 degrees Celsius. Recent advances in '
        'nano-filler technology suggest that incorporating ceramic nanoparticles into the polymer '
        'matrix can enhance thermal stability without compromising processability.'
    )
    doc.add_paragraph(
        'The primary objective of this study was to evaluate the thermal behavior of three distinct '
        'CFRP formulations: (a) baseline epoxy/CF, (b) epoxy/CF with 1.0 wt% nano-silica, and '
        '(c) epoxy/CF with 2.5 wt% nano-silica. Each formulation was characterized using DSC, TGA, '
        'and dynamic mechanical analysis (DMA) to establish a complete thermal profile.'
    )
    doc.add_paragraph(
        'Previous work by Zhang et al. (2023) demonstrated that nano-silica loadings above 3.0 wt% '
        'led to agglomeration and reduced fracture toughness. Therefore, our study focused on the '
        'sub-3.0 wt% regime to optimize the balance between thermal enhancement and mechanical '
        'performance. The results of this investigation have implications for next-generation '
        'composite design in high-temperature structural applications.'
    )

    # ===== PAGE 3: Experimental Setup (with Image 1) =====
    add_page_break(doc)
    doc.add_heading('2. Experimental Methods', level=1)
    doc.add_heading('2.1 Materials and Sample Preparation', level=2)
    doc.add_paragraph(
        'High-modulus carbon fiber (Toray T700S, 12K tow) was used as reinforcement. The matrix '
        'system consisted of bisphenol-A diglycidyl ether (DGEBA) epoxy resin cured with '
        'diethylenetriamine (DETA) hardener at a stoichiometric ratio. Nano-silica particles '
        '(Aerosil R972, average diameter 16 nm) were pre-dispersed in the resin using a high-shear '
        'mixer operating at 3000 RPM for 45 minutes, followed by ultrasonication for 30 minutes '
        'to minimize agglomeration.'
    )
    doc.add_paragraph(
        'Composite panels (300 x 300 x 3 mm) were fabricated using vacuum-assisted resin transfer '
        'molding (VARTM). The fiber volume fraction was maintained at 55 +/- 2% across all samples. '
        'Curing was performed at 120 degrees Celsius for 4 hours, followed by post-cure at '
        '180 degrees Celsius for 2 hours. Test specimens were cut using a water-cooled diamond saw '
        'to minimize thermal damage during machining.'
    )

    doc.add_heading('2.2 Instrumentation Setup', level=2)
    doc.add_paragraph(
        'The experimental apparatus configuration is shown below. All thermal measurements were '
        'conducted in a nitrogen atmosphere to prevent oxidative degradation during testing.'
    )

    # Image 1: Experimental Setup
    doc.add_picture(image_paths[0], width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        'Temperature calibration was performed using indium and zinc standards prior to each '
        'measurement session. The DSC cell constant was verified to be within 1.5% of the '
        'theoretical value for indium melting enthalpy (28.45 J/g).'
    )

    # ===== PAGE 4-5: More Methods =====
    add_page_break(doc)
    doc.add_heading('2.3 Differential Scanning Calorimetry (DSC)', level=2)
    doc.add_paragraph(
        'DSC measurements were performed on a TA Instruments Q2000 equipped with an RCS90 cooling '
        'system. Samples (8-12 mg) were sealed in hermetic aluminum pans. The heating protocol '
        'consisted of three segments: (1) equilibrate at -50 degrees Celsius for 5 minutes, '
        '(2) heat at 10 degrees Celsius/min to 300 degrees Celsius, (3) cool at 10 degrees Celsius/min '
        'to -50 degrees Celsius. The second heating scan was used for Tg determination to eliminate '
        'thermal history effects.'
    )
    doc.add_paragraph(
        'Glass transition temperatures were determined using the midpoint method (ASTM E1356). '
        'Three replicate measurements were performed for each formulation, and mean values with '
        'standard deviations are reported. The onset, midpoint, and endset temperatures were '
        'recorded for each transition.'
    )

    doc.add_heading('2.4 Thermogravimetric Analysis (TGA)', level=2)
    doc.add_paragraph(
        'TGA was conducted on a TA Instruments Q500 under nitrogen flow (60 mL/min). Samples '
        '(15-20 mg) were heated from 30 to 800 degrees Celsius at 10 degrees Celsius/min. The '
        'onset of decomposition was defined as the temperature at 5% mass loss (T5%). The '
        'derivative thermogravimetric (DTG) curve was used to identify the temperature of maximum '
        'mass loss rate (Tmax).'
    )
    doc.add_paragraph(
        'Isothermal TGA experiments were also performed at 200, 250, and 300 degrees Celsius for '
        '2 hours each to evaluate long-term thermal stability relevant to aerospace service '
        'conditions. The activation energy of decomposition was calculated using the Flynn-Wall-Ozawa '
        'method from TGA data collected at heating rates of 5, 10, 15, and 20 degrees Celsius/min.'
    )

    doc.add_heading('2.5 Dynamic Mechanical Analysis (DMA)', level=2)
    doc.add_paragraph(
        'DMA was performed in three-point bending mode on a TA Instruments Q800. Rectangular '
        'specimens (60 x 12 x 3 mm) were tested at 1 Hz frequency with an oscillation amplitude '
        'of 15 micrometers. Temperature sweeps were conducted from -50 to 250 degrees Celsius at '
        '3 degrees Celsius/min. Storage modulus (E\'), loss modulus (E\'\'), and tan delta were recorded.'
    )

    # ===== PAGE 5-6: Results Section with Image 2 =====
    add_page_break(doc)
    doc.add_heading('3. Results and Discussion', level=1)
    doc.add_heading('3.1 DSC Results - Glass Transition Behavior', level=2)
    doc.add_paragraph(
        'The DSC results revealed a clear trend of increasing glass transition temperature with '
        'nano-silica content. The baseline CFRP exhibited a Tg of 142.7 +/- 1.2 degrees Celsius '
        '(midpoint), while the 1.0 wt% and 2.5 wt% nano-silica formulations showed Tg values of '
        '151.3 +/- 0.8 and 161.0 +/- 1.5 degrees Celsius, respectively. This represents an '
        'improvement of 8.6K and 18.3K over the baseline.'
    )

    # Table of DSC results
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['Sample ID', 'Nano-silica (wt%)', 'Tg Onset (C)', 'Tg Midpoint (C)', 'Delta Cp (J/g-C)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data = [
        ['CFRP-B1', '0.0', '138.2', '142.7', '0.342'],
        ['CFRP-N1', '1.0', '146.8', '151.3', '0.318'],
        ['CFRP-N2', '2.5', '155.4', '161.0', '0.295'],
        ['CFRP-N3*', '3.5', '157.1', '163.2', '0.271'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph(
        '* CFRP-N3 included for reference but excluded from primary analysis due to observed '
        'nanoparticle agglomeration.'
    )
    doc.add_paragraph(
        'The reduction in heat capacity change (Delta Cp) with increasing nano-silica content '
        'suggests restricted segmental mobility of the polymer chains near the filler-matrix '
        'interface. This constrained layer effect has been documented in similar nanocomposite '
        'systems and is consistent with the observed Tg elevation.'
    )

    doc.add_paragraph(
        'The distribution of raw measurement data across all thermal analysis runs is presented below.'
    )

    # Image 2: Raw Data Distribution
    doc.add_picture(image_paths[1], width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ===== PAGE 7-8: More Results =====
    add_page_break(doc)
    doc.add_heading('3.2 TGA Results - Thermal Decomposition', level=2)
    doc.add_paragraph(
        'Thermogravimetric analysis confirmed enhanced thermal stability with nano-silica '
        'incorporation. The baseline CFRP showed T5% at 327.4 degrees Celsius with Tmax at '
        '381.2 degrees Celsius. The 2.5 wt% nano-silica formulation exhibited T5% at 342.8 '
        'degrees Celsius (+15.4K) and Tmax at 396.7 degrees Celsius (+15.5K). Char yield at '
        '800 degrees Celsius increased from 42.3% (baseline) to 48.1% (2.5 wt% nano-silica), '
        'indicating improved flame retardancy potential.'
    )

    table2 = doc.add_table(rows=4, cols=5)
    table2.style = 'Table Grid'
    headers2 = ['Sample', 'T5% (C)', 'Tmax (C)', 'Char Yield (%)', 'Ea (kJ/mol)']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data2 = [
        ['CFRP-B1', '327.4', '381.2', '42.3', '168.5'],
        ['CFRP-N1', '335.1', '389.4', '45.7', '175.2'],
        ['CFRP-N2', '342.8', '396.7', '48.1', '183.9'],
    ]
    for r, row_data in enumerate(data2, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_paragraph(
        'The Flynn-Wall-Ozawa analysis yielded activation energies of 168.5, 175.2, and 183.9 kJ/mol '
        'for the baseline, 1.0 wt%, and 2.5 wt% formulations respectively. The increasing activation '
        'energy trend confirms that the nano-silica particles create a physical barrier that impedes '
        'the diffusion of volatile decomposition products from the polymer matrix.'
    )

    doc.add_paragraph(
        'Isothermal TGA at 250 degrees Celsius for 2 hours showed mass losses of 2.8%, 1.9%, '
        'and 1.2% for the baseline, 1.0 wt%, and 2.5 wt% formulations. These results are '
        'particularly significant for aerospace applications where sustained elevated temperature '
        'exposure is expected during service life.'
    )

    add_page_break(doc)
    doc.add_heading('3.3 DMA Results - Viscoelastic Behavior', level=2)
    doc.add_paragraph(
        'Dynamic mechanical analysis provided additional insight into the thermo-mechanical '
        'response. The storage modulus at 25 degrees Celsius was 68.3, 72.1, and 76.8 GPa for '
        'the baseline, 1.0 wt%, and 2.5 wt% nano-silica formulations. The tan delta peak '
        'temperature, another measure of Tg, followed the same trend: 148.6, 157.2, and 167.4 '
        'degrees Celsius, respectively.'
    )
    doc.add_paragraph(
        'The breadth of the tan delta peak increased with nano-silica content, indicating a '
        'broader distribution of relaxation times. This broadening is attributed to the heterogeneous '
        'constraint environment created by the nanoparticles, where polymer chains near the filler '
        'surface have significantly restricted mobility compared to those in the bulk matrix.'
    )
    doc.add_paragraph(
        'The storage modulus retention ratio (E\' at 200C / E\' at 25C) improved from 0.31 (baseline) '
        'to 0.44 (2.5 wt% nano-silica), demonstrating superior high-temperature stiffness. This '
        'finding has direct implications for structural applications where elevated temperature '
        'creep resistance is critical.'
    )

    # ===== PAGE 9: Analysis Results (with Image 3) =====
    add_page_break(doc)
    doc.add_heading('3.4 Statistical Analysis of Combined Results', level=2)
    doc.add_paragraph(
        'A comprehensive statistical analysis was performed using one-way ANOVA with Tukey\'s '
        'post-hoc comparison. All pairwise comparisons between formulations for Tg, T5%, and '
        'storage modulus were statistically significant (p < 0.01). The correlation between '
        'nano-silica content and Tg improvement was well-described by a linear model '
        '(R-squared = 0.994), suggesting that the enhancement mechanism is consistent across '
        'the studied concentration range.'
    )
    doc.add_paragraph(
        'The compiled analysis results showing the correlation between filler content and thermal '
        'properties are presented below.'
    )

    # Image 3: Analysis Results
    doc.add_picture(image_paths[2], width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        'Regression analysis indicates that each 1.0 wt% increase in nano-silica content '
        'yields approximately 7.3K improvement in Tg and 6.2K improvement in T5%. However, '
        'the diminishing returns observed at higher loadings (CFRP-N3 data point) suggest '
        'an optimal loading range of 2.0-3.0 wt% for this specific resin system.'
    )

    # ===== PAGE 10-11: More Discussion =====
    add_page_break(doc)
    doc.add_heading('3.5 Microstructural Analysis', level=2)
    doc.add_paragraph(
        'Scanning electron microscopy (SEM) of fracture surfaces revealed a progressive change '
        'in failure morphology with increasing nano-silica content. The baseline CFRP exhibited '
        'relatively smooth, mirror-like fracture surfaces characteristic of brittle matrix failure. '
        'In contrast, the 2.5 wt% nano-silica formulation showed significantly rougher surfaces '
        'with evidence of crack deflection, particle pull-out, and localized plastic deformation '
        'of the matrix.'
    )
    doc.add_paragraph(
        'Energy-dispersive X-ray spectroscopy (EDS) mapping confirmed uniform dispersion of '
        'silicon across the fracture surface for the 1.0 and 2.5 wt% formulations. No silica-rich '
        'domains exceeding 200 nm were observed, confirming effective dispersion achieved during '
        'the mixing protocol. For the 3.5 wt% reference sample (CFRP-N3), occasional agglomerates '
        'up to 1.5 micrometers were detected, consistent with the observed plateau in Tg improvement.'
    )
    doc.add_paragraph(
        'Transmission electron microscopy (TEM) of ultramicrotomed thin sections revealed an '
        'interphase region of approximately 5-8 nm surrounding each nanoparticle, where the '
        'polymer chain packing density appeared to differ from the bulk. This interphase region '
        'is believed to be the primary contributor to the constrained layer effect observed in '
        'both DSC and DMA measurements.'
    )

    add_page_break(doc)
    doc.add_heading('3.6 Thermal Cycling Performance', level=2)
    doc.add_paragraph(
        'Accelerated thermal cycling tests (-40 to 180 degrees Celsius, 1000 cycles) were '
        'performed to assess long-term durability. Post-cycling DSC measurements showed Tg '
        'retention of 96.2%, 97.8%, and 98.5% for the baseline, 1.0 wt%, and 2.5 wt% '
        'formulations. Ultrasonic C-scan imaging revealed no significant delamination in any '
        'sample, although baseline specimens showed slight increases in attenuation coefficient '
        '(+0.3 dB/mm) suggesting micro-crack development in the matrix-rich regions.'
    )
    doc.add_paragraph(
        'Flexural strength retention after thermal cycling was 91.4%, 94.7%, and 96.2% for the '
        'three formulations. The nano-silica particles appear to act as crack arrest sites, '
        'reducing the propagation of thermally-induced micro-cracks. This crack-bridging mechanism '
        'is consistent with previous observations in nano-filled polymer systems and contributes '
        'to the overall improvement in thermal fatigue resistance.'
    )

    # ===== PAGE 12: Comparison (with Image 4) =====
    add_page_break(doc)
    doc.add_heading('4. Comparative Discussion', level=1)
    doc.add_paragraph(
        'To contextualize our findings, a comparison with published results for similar '
        'nanocomposite systems was performed. Our observed Tg improvement of 7.3K per wt% '
        'nano-silica is in excellent agreement with values reported by Chen et al. (2024) for '
        'DGEBA/nano-silica systems (7.8K/wt%) and slightly higher than the 5.9K/wt% reported '
        'by Muller et al. (2023) for a different epoxy chemistry. The higher enhancement in our '
        'system may be attributed to the smaller particle size (16 nm vs. 25 nm) and improved '
        'dispersion achieved through the combined high-shear/ultrasonication protocol.'
    )
    doc.add_paragraph(
        'The comparative analysis across all formulations is visualized below.'
    )

    # Image 4: Comparison Chart
    doc.add_picture(image_paths[3], width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        'The comparison demonstrates that our 2.5 wt% formulation achieves thermal performance '
        'competitive with systems requiring higher filler loadings, while maintaining superior '
        'processability as evidenced by the viscosity measurements presented in Appendix A.'
    )

    # ===== PAGE 13: Conclusions =====
    add_page_break(doc)
    doc.add_heading('5. Conclusions', level=1)
    doc.add_paragraph(
        'This study systematically evaluated the effect of nano-silica incorporation on the '
        'thermal properties of carbon-fiber reinforced polymer composites. The key findings are:'
    )
    conclusions = [
        'The glass transition temperature increased linearly with nano-silica content at a rate '
        'of 7.3K per wt%, reaching 161.0 degrees Celsius at 2.5 wt% loading (18.3K improvement '
        'over baseline).',
        'Thermal decomposition onset (T5%) improved by 15.4K with 2.5 wt% nano-silica, and '
        'char yield increased by 5.8 percentage points, indicating enhanced thermal stability.',
        'Dynamic mechanical analysis confirmed a 12.4% increase in room-temperature storage '
        'modulus and significantly improved high-temperature stiffness retention (0.44 vs. 0.31).',
        'Thermal cycling tests demonstrated superior fatigue resistance with nano-silica '
        'incorporation, attributed to crack-bridging and energy dissipation mechanisms.',
        'An optimal nano-silica loading of 2.0-3.0 wt% is recommended to maximize thermal '
        'enhancement while avoiding agglomeration-related degradation.',
    ]
    for item in conclusions:
        doc.add_paragraph(item, style='List Number')

    # ===== PAGE 14: References =====
    add_page_break(doc)
    doc.add_heading('6. References', level=1)
    references = [
        'Chen, W., Liu, Y., & Park, S. (2024). "Nano-silica modified epoxy composites: Thermal '
        'and mechanical characterization." Composites Science and Technology, 245, 110328.',
        'Muller, F., Schmidt, K., & Weber, H. (2023). "Effect of nanoparticle size on thermal '
        'properties of filled epoxy systems." Polymer Engineering & Science, 63(8), 1842-1855.',
        'Zhang, L., Wang, R., & Tanaka, M. (2023). "Agglomeration effects in high-loading '
        'nanocomposites: A critical review." Materials Today Advances, 18, 100375.',
        'ASTM E1356-08. (2014). "Standard Test Method for Assignment of the Glass Transition '
        'Temperatures by Differential Scanning Calorimetry." ASTM International.',
        'Flynn, J.H. & Wall, L.A. (1966). "A quick, direct method for the determination of '
        'activation energy from thermogravimetric data." Journal of Polymer Science Part B, '
        '4(5), 323-328.',
        'Wetzel, B., Haupert, F., & Friedrich, K. (2021). "Influence of surface treatment on '
        'nano-silica/epoxy composite properties." Composites Part A, 142, 106215.',
        'Park, J.H., & Jana, S.C. (2022). "Mechanism of exfoliation of nanoclay in epoxy-clay '
        'nanocomposites." Macromolecules, 55(14), 5987-5999.',
        'Johnsen, B.B., Kinloch, A.J., & Taylor, A.C. (2023). "Toughness of syndiotactic '
        'polystyrene/epoxy polymer blends." Polymer, 46(20), 7352-7369.',
    ]
    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f'[{i}] {ref}')

    # ===== PAGE 15: List of Figures (empty - to be filled by agent) =====
    add_page_break(doc)
    doc.add_heading('List of Figures', level=1)
    # No content below - the agent needs to insert a Table of Figures here

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
