"""
Reward Script: Set outer table border to 2.5pt solid dark blue, remove all inner borders
Task ID: writer_tbl_029
Domain: libreoffice_writer
Scoring:
  Component 1 (0.40): Table outer borders are 2.5pt (sz=20) single dark blue (color~00008B)
  Component 2 (0.30): Table inner borders (insideH/insideV) are removed (val=none)
  Component 3 (0.30): Cell contents unchanged (4 rows x 3 cols with expected text)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_029'
FILE_PATH = f'{WORKDIR}/Desktop/featured_table.docx'

# Expected cell contents (row, col) -> text
EXPECTED_CELLS = {
    (0, 0): 'Feature',
    (0, 1): 'Basic',
    (0, 2): 'Premium',
    (1, 0): 'Storage',
    (1, 1): '5 GB',
    (1, 2): '100 GB',
    (2, 0): 'Support',
    (2, 1): 'Email',
    (2, 2): '24/7 Phone',
    (3, 0): 'Price',
    (3, 1): '$5/mo',
    (3, 2): '$25/mo',
}

# Dark blue color variants accepted (00008B = dark blue)
DARK_BLUE_COLORS = {'00008B', '000080', '00008b'}

def get_border_attr(border_elem, attr):
    """Get attribute from border element, return None if not present."""
    if border_elem is None:
        return None
    return border_elem.get(qn(f'w:{attr}'))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must load
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least one table
    if not doc.tables:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    tbl = table._tbl

    # ----------------------------------------------------------------
    # Component 1: Outer borders are 2.5pt (sz=20) single dark blue (0.40 points)
    # Task change: initial has sz=4, black; golden has sz=20, dark blue (00008B)
    # ----------------------------------------------------------------
    try:
        tblPr = tbl.find(qn('w:tblPr'))
        outer_pass = False
        outer_details = []

        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                outer_sides = ['top', 'left', 'bottom', 'right']
                all_outer_correct = True
                for side in outer_sides:
                    elem = tblBorders.find(qn(f'w:{side}'))
                    if elem is None:
                        all_outer_correct = False
                        outer_details.append(f"{side}: missing")
                        continue
                    val = elem.get(qn('w:val'))
                    sz = elem.get(qn('w:sz'))
                    color = elem.get(qn('w:color'), '').upper()
                    # Accept sz=20 (2.5pt = 20 eighths-of-a-point) and dark blue color
                    # sz tolerance: allow 18-22 range for minor rounding
                    sz_ok = sz is not None and 18 <= int(sz) <= 22
                    color_ok = color.lstrip('#') in {c.upper() for c in DARK_BLUE_COLORS}
                    val_ok = val == 'single'
                    if not (sz_ok and color_ok and val_ok):
                        all_outer_correct = False
                        outer_details.append(
                            f"{side}: val={val}, sz={sz}, color={color} (expected single/20/00008B)"
                        )
                    else:
                        outer_details.append(f"{side}: OK (val={val}, sz={sz}, color={color})")
                outer_pass = all_outer_correct
            else:
                outer_details.append("No tblBorders element found")
        else:
            outer_details.append("No tblPr element found")

        if outer_pass:
            print(f"PASS: Component 1 — All outer borders are 2.5pt single dark blue (0.40 pts)")
            for d in outer_details:
                print(f"  {d}")
            total_score += 0.40
        else:
            print(f"FAIL: Component 1 — Outer borders not correct (expected sz=20, color=00008B, val=single)")
            for d in outer_details:
                print(f"  {d}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ----------------------------------------------------------------
    # Component 2: Inner borders (insideH, insideV) are removed/none (0.30 points)
    # Task change: initial has insideH/V as single; golden has insideH/V as none
    # ----------------------------------------------------------------
    try:
        tblPr = tbl.find(qn('w:tblPr'))
        inner_pass = False
        inner_details = []

        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                inner_sides = ['insideH', 'insideV']
                all_inner_none = True
                for side in inner_sides:
                    elem = tblBorders.find(qn(f'w:{side}'))
                    if elem is None:
                        # Missing element can also mean no border (acceptable)
                        inner_details.append(f"{side}: missing (treated as none)")
                        continue
                    val = elem.get(qn('w:val'))
                    if val not in ('none', 'nil', None):
                        all_inner_none = False
                        sz = elem.get(qn('w:sz'))
                        color = elem.get(qn('w:color'))
                        inner_details.append(
                            f"{side}: val={val}, sz={sz}, color={color} (expected none)"
                        )
                    else:
                        inner_details.append(f"{side}: OK (val={val})")
                inner_pass = all_inner_none
            else:
                inner_details.append("No tblBorders element — inner borders implicitly present")
                inner_pass = False
        else:
            inner_details.append("No tblPr element")
            inner_pass = False

        if inner_pass:
            print(f"PASS: Component 2 — Inner borders (insideH/insideV) are removed (0.30 pts)")
            for d in inner_details:
                print(f"  {d}")
            total_score += 0.30
        else:
            print(f"FAIL: Component 2 — Inner borders not removed")
            for d in inner_details:
                print(f"  {d}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ----------------------------------------------------------------
    # Component 3: Cell contents unchanged (4 rows x 3 cols, expected text) (0.30 points)
    # The task only modifies borders, not content.
    # This check FAILS on initial if we use the outer borders as a gate:
    #   We verify content ONLY after confirming outer borders changed (via comp 1 pass).
    # BUT to avoid scoring pre-existing properties, we tie this to the border changes:
    #   Content preservation is only awarded if the table ALSO has dark blue outer borders
    #   (i.e., the task was actually completed and content was not damaged).
    # So this component is conditional on comp 1 passing.
    # ----------------------------------------------------------------
    try:
        if outer_pass:
            rows = len(table.rows)
            cols = len(table.columns)
            dims_ok = (rows == 4 and cols == 3)
            content_errors = []

            for (ri, ci), expected in EXPECTED_CELLS.items():
                actual = table.cell(ri, ci).text.strip()
                if actual != expected:
                    content_errors.append(f"Cell[{ri}][{ci}]: expected '{expected}', got '{actual}'")

            if dims_ok and not content_errors:
                print(f"PASS: Component 3 — Cell contents preserved (4x3 table, all text correct) (0.30 pts)")
                total_score += 0.30
            else:
                if not dims_ok:
                    print(f"FAIL: Component 3 — Table dimensions wrong: {rows}x{cols} (expected 4x3)")
                else:
                    print(f"FAIL: Component 3 — Cell content errors:")
                    for err in content_errors:
                        print(f"  {err}")
        else:
            # Outer borders not correct — comp 3 purposely not awarded
            # (skipping content check since outer border change is required for this component)
            print("SKIP: Component 3 — skipped (outer border change not detected, initial env guard)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
