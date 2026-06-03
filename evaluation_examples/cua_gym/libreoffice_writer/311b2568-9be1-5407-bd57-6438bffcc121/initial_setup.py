"""
Initial Setup: History essay document with ordinal indicators as regular text
Task ID: writer_txtfmt_060
Domain: libreoffice_writer

Creates a .docx file at /home/user/Desktop/history_essay.docx with the
sentence containing '1st', '2nd', '3rd', '4th' — ordinal suffixes are
plain text (NOT superscript). The agent must format them as superscript.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_txtfmt_060'
OUTPUT = f'{WORKDIR}/history_essay.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure the Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a title paragraph
    title = doc.add_paragraph()
    title_run = title.add_run("The Continental Congress and Early American Resistance")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)

    # Add an introductory paragraph
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "The foundations of American independence were laid through a series of pivotal "
        "meetings and acts of collective defiance against British rule. Colonial leaders "
        "recognized the need for unified action to address grievances and protect their "
        "rights as British subjects."
    )
    intro_run.font.name = "Times New Roman"
    intro_run.font.size = Pt(12)

    # The key paragraph with ordinal indicators — all as plain text (no superscript)
    # Build the paragraph run-by-run to precisely control each piece of text
    key_para = doc.add_paragraph()

    # "The "
    r1 = key_para.add_run("The ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)

    # "1" (number — normal)
    r2 = key_para.add_run("1")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

    # "st" — plain text (NOT superscript in initial state)
    r3 = key_para.add_run("st")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)
    r3.font.superscript = False

    # " Continental Congress convened in 1774, and the "
    r4 = key_para.add_run(" Continental Congress convened in 1774, and the ")
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(12)

    # "2" (number — normal)
    r5 = key_para.add_run("2")
    r5.font.name = "Times New Roman"
    r5.font.size = Pt(12)

    # "nd" — plain text (NOT superscript in initial state)
    r6 = key_para.add_run("nd")
    r6.font.name = "Times New Roman"
    r6.font.size = Pt(12)
    r6.font.superscript = False

    # " followed in 1775. By the "
    r7 = key_para.add_run(" followed in 1775. By the ")
    r7.font.name = "Times New Roman"
    r7.font.size = Pt(12)

    # "3" (number — normal)
    r8 = key_para.add_run("3")
    r8.font.name = "Times New Roman"
    r8.font.size = Pt(12)

    # "rd" — plain text (NOT superscript in initial state)
    r9 = key_para.add_run("rd")
    r9.font.name = "Times New Roman"
    r9.font.size = Pt(12)
    r9.font.superscript = False

    # " year of conflict, and into the "
    r10 = key_para.add_run(" year of conflict, and into the ")
    r10.font.name = "Times New Roman"
    r10.font.size = Pt(12)

    # "4" (number — normal)
    r11 = key_para.add_run("4")
    r11.font.name = "Times New Roman"
    r11.font.size = Pt(12)

    # "th" — plain text (NOT superscript in initial state)
    r12 = key_para.add_run("th")
    r12.font.name = "Times New Roman"
    r12.font.size = Pt(12)
    r12.font.superscript = False

    # ", the colonies had organized significant resistance."
    r13 = key_para.add_run(", the colonies had organized significant resistance.")
    r13.font.name = "Times New Roman"
    r13.font.size = Pt(12)

    # Add a closing paragraph with additional context
    closing = doc.add_paragraph()
    closing_run = closing.add_run(
        "The delegates who gathered represented a cross-section of colonial society: "
        "merchants, lawyers, planters, and physicians. Their deliberations would shape "
        "the political philosophy of a new nation and establish principles of governance "
        "that endure to this day. The courage demonstrated during these formative years "
        "remains a testament to the power of principled resistance against tyranny."
    )
    closing_run.font.name = "Times New Roman"
    closing_run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
