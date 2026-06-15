"""
Reward Script: Resize table columns (5cm, 7cm, 4cm) and merge last row cells into footer
Task ID: osworld_writer_table_editing_004
Domain: libreoffice_writer
Scoring:
  - Component 1: Column widths match task spec (5cm, 7cm, 4cm) — 0.6 points
  - Component 2: Last row cells merged spanning all 3 columns — 0.4 points
  Total: 1.0
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_table_editing_004'

# Tolerance for column width comparison: +/- 36 dxa (~0.635 cm)
# 1 cm = 567.0 dxa (1440 / 2.54 = 566.93...)
CM_TO_DXA = 1440 / 2.54
COL_WIDTH_TOLERANCE_DXA = 72  # ~1.27 mm — generous for rounding differences

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: file must have at least one table with 5 rows and 3 columns
    if not doc.tables:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    if len(table.rows) < 2:
        print(f"FAIL: Table has fewer than 2 rows (found {len(table.rows)})")
        print("REWARD: 0.0")
        return 0.0

    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # -------------------------------------------------------------------------
    # Component 1: Column widths match task spec — 5cm, 7cm, 4cm (0.6 points)
    # Check first non-merged row (row 0) for column widths.
    # Initial file has all cols at 2880 dxa (~5.08cm each); golden has:
    #   Col0=2835 dxa (~5.0cm), Col1=3969 dxa (~7.0cm), Col2=2268 dxa (~4.0cm)
    # -------------------------------------------------------------------------
    try:
        target_widths_cm = [5.0, 7.0, 4.0]
        target_widths_dxa = [w * CM_TO_DXA for w in target_widths_cm]

        # Read widths from the first row (header row, definitely not merged)
        first_row = table.rows[0]
        actual_widths_dxa = []
        for cell in first_row.cells:
            tcPr = cell._tc.tcPr
            if tcPr is not None:
                tcW = tcPr.find(f'{{{NS}}}tcW')
                if tcW is not None:
                    w_str = tcW.get(f'{{{NS}}}w')
                    if w_str is not None:
                        actual_widths_dxa.append(int(w_str))
                        continue
            actual_widths_dxa.append(None)

        if len(actual_widths_dxa) < 3:
            print(f"FAIL: Component 1 — could not read 3 column widths from row 0, got: {actual_widths_dxa}")
        else:
            all_match = True
            for idx, (actual, target) in enumerate(zip(actual_widths_dxa[:3], target_widths_dxa)):
                if actual is None:
                    print(f"FAIL: Component 1 — Col{idx} width is None")
                    all_match = False
                    break
                diff = abs(actual - target)
                actual_cm = actual / CM_TO_DXA
                target_cm = target_widths_cm[idx]
                if diff > COL_WIDTH_TOLERANCE_DXA:
                    print(f"FAIL: Component 1 — Col{idx} expected ~{target_cm:.1f}cm ({target:.0f} dxa), "
                          f"found {actual_cm:.4f}cm ({actual} dxa), diff={diff:.1f} dxa")
                    all_match = False
                else:
                    print(f"  Col{idx}: {actual_cm:.4f}cm ({actual} dxa) ≈ {target_cm:.1f}cm — OK")

            if all_match:
                print(f"PASS: Component 1 — Column widths match spec "
                      f"(Col0={actual_widths_dxa[0]} dxa, Col1={actual_widths_dxa[1]} dxa, "
                      f"Col2={actual_widths_dxa[2]} dxa) (0.6 pts)")
                total_score += 0.6
    except Exception as e:
        print(f"ERROR: Component 1 — could not check column widths: {e}")

    # -------------------------------------------------------------------------
    # Component 2: Last row cells merged spanning all 3 columns (0.4 points)
    # In the golden file the last row has exactly 1 <w:tc> element with
    # gridSpan=3. In the initial file the last row has 3 separate <w:tc> cells.
    # -------------------------------------------------------------------------
    try:
        last_row = table.rows[-1]

        # Count actual <w:tc> elements in the last row's XML
        # (python-docx expands merged cells so len(last_row.cells) is unreliable)
        tc_elements = last_row._tr.findall(f'{{{NS}}}tc')
        num_tcs = len(tc_elements)

        if num_tcs == 0:
            print("FAIL: Component 2 — no <w:tc> elements found in last row")
        elif num_tcs == 1:
            # Exactly one cell — now verify it spans all 3 columns via gridSpan
            tc = tc_elements[0]
            tcPr = tc.find(f'{{{NS}}}tcPr')
            gridSpan_val = None
            if tcPr is not None:
                gs = tcPr.find(f'{{{NS}}}gridSpan')
                if gs is not None:
                    gridSpan_val = gs.get(f'{{{NS}}}val')

            if gridSpan_val is not None and int(gridSpan_val) >= 3:
                print(f"PASS: Component 2 — Last row is merged into single cell "
                      f"(gridSpan={gridSpan_val}) (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — Last row has 1 cell but gridSpan={gridSpan_val} "
                      f"(expected 3)")
        else:
            # More than 1 cell — not fully merged
            print(f"FAIL: Component 2 — Last row has {num_tcs} separate <w:tc> elements "
                  f"(expected 1 merged cell with gridSpan=3)")
    except Exception as e:
        print(f"ERROR: Component 2 — could not check last row merge: {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
