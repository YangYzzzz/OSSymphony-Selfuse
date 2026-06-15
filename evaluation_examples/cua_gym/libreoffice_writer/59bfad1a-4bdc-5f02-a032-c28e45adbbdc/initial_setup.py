"""
Initial Setup: Product table in LibreOffice Writer with 5 rows and 3 columns (equal default widths)
Task ID: osworld_writer_table_editing_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_table_editing_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title heading
    heading = doc.add_heading('Product Catalog', level=1)

    # Add brief intro paragraph
    intro = doc.add_paragraph(
        'The following table lists available products with their specifications and pricing details. '
        'Please review the catalog and update column widths as needed for better readability.'
    )

    # Create a 5-row, 3-column product table with default equal column widths
    # Rows: 1 header + 4 data rows = 5 rows total
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # ---- Row 0: Header ----
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product Name'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Unit Price'

    # Make header row bold
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # ---- Row 1 ----
    row1 = table.rows[1].cells
    row1[0].text = 'Alpine Trekking Boots'
    row1[1].text = 'Waterproof leather boots with ankle support, suitable for mountain trails'
    row1[2].text = '$189.99'

    # ---- Row 2 ----
    row2 = table.rows[2].cells
    row2[0].text = 'TrailMaster Backpack'
    row2[1].text = '65L capacity with ergonomic harness system and hydration sleeve'
    row2[2].text = '$124.50'

    # ---- Row 3 ----
    row3 = table.rows[3].cells
    row3[0].text = 'Summit Sleeping Bag'
    row3[1].text = 'Rated -10°C, lightweight 850-fill down insulation, compression sack included'
    row3[2].text = '$249.00'

    # ---- Row 4 (last row — no merge in initial state) ----
    row4 = table.rows[4].cells
    row4[0].text = 'Navigation Compass'
    row4[1].text = 'Precision baseplate compass with declination adjustment and magnifier'
    row4[2].text = '$42.75'

    # NOTE: Column widths are left at default equal widths (NOT resized)
    # NOTE: Last row cells are NOT merged (that is the task to do)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
