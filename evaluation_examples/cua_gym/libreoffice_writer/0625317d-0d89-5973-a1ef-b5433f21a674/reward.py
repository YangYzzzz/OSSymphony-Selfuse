"""
Reward Script: Set multi-language spell-checking in Writer document
Task ID: writer_fp_032
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Body text (P0-P5) language set to en-US
  Component 2 (0.35): Appendix A (P6-P9) language set to de-DE
  Component 3 (0.30): Appendix B (P10-P13) language set to es-ES
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_032'


def get_run_language(run):
    """Extract w:lang w:val from a run's rPr element. Returns None if not set."""
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        return None
    lang = rpr.find(qn('w:lang'))
    if lang is None:
        return None
    return lang.get(qn('w:val'))


def check_paragraphs_language(doc, start_idx, end_idx, expected_lang):
    """
    Check that all runs in paragraphs[start_idx:end_idx+1] have the expected language.
    Returns (passes, total_runs) where passes is number of runs with correct lang.
    Paragraphs with no runs (empty) are skipped.
    """
    total_runs = 0
    matching_runs = 0
    for i in range(start_idx, end_idx + 1):
        if i >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[i]
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            lang = get_run_language(run)
            if lang == expected_lang:
                matching_runs += 1
    return matching_runs, total_runs


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_paras = len(doc.paragraphs)
    print(f"INFO: Document has {num_paras} paragraphs")

    # Identify section boundaries by looking for "Appendix A" and "Appendix B" headings
    appendix_a_start = None
    appendix_b_start = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "Appendix A":
            appendix_a_start = i
        elif text == "Appendix B":
            appendix_b_start = i

    if appendix_a_start is None or appendix_b_start is None:
        print(f"CRITICAL: Cannot find section headings. Appendix A at {appendix_a_start}, Appendix B at {appendix_b_start}")
        print("REWARD: 0.0")
        return 0.0

    # Section ranges:
    # Body: P0 to (appendix_a_start - 1)
    # Appendix A: appendix_a_start to (appendix_b_start - 1)
    # Appendix B: appendix_b_start to end
    body_end = appendix_a_start - 1
    app_a_end = appendix_b_start - 1
    app_b_end = num_paras - 1

    print(f"INFO: Body P0-P{body_end}, Appendix A P{appendix_a_start}-P{app_a_end}, Appendix B P{appendix_b_start}-P{app_b_end}")

    # Component 1: Body text has language en-US (0.35 points)
    try:
        matching, total = check_paragraphs_language(doc, 0, body_end, 'en-US')
        if total == 0:
            print(f"FAIL: Component 1 — No runs found in body text")
        elif matching == total:
            print(f"PASS: Component 1 — All {total} body runs have lang=en-US (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — {matching}/{total} body runs have lang=en-US")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Appendix A has language de-DE (0.35 points)
    try:
        matching, total = check_paragraphs_language(doc, appendix_a_start, app_a_end, 'de-DE')
        if total == 0:
            print(f"FAIL: Component 2 — No runs found in Appendix A")
        elif matching == total:
            print(f"PASS: Component 2 — All {total} Appendix A runs have lang=de-DE (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 2 — {matching}/{total} Appendix A runs have lang=de-DE")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Appendix B has language es-ES (0.30 points)
    try:
        matching, total = check_paragraphs_language(doc, appendix_b_start, app_b_end, 'es-ES')
        if total == 0:
            print(f"FAIL: Component 3 — No runs found in Appendix B")
        elif matching == total:
            print(f"PASS: Component 3 — All {total} Appendix B runs have lang=es-ES (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 3 — {matching}/{total} Appendix B runs have lang=es-ES")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
