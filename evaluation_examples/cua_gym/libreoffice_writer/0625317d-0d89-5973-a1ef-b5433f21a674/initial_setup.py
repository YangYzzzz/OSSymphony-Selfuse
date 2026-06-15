"""
Initial Setup: Multi-language spell-check document
Task ID: writer_fp_032
Domain: libreoffice_writer

Creates a Writer document with English body text, German Appendix A,
and Spanish Appendix B. All text has NO explicit language set (default).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # =============================================
    # MAIN BODY (English)
    # =============================================
    heading = doc.add_heading('Quarterly Business Performance Report', level=1)

    doc.add_paragraph(
        'The third quarter of fiscal year 2025 has demonstrated remarkable growth '
        'across all major business divisions. Revenue increased by 18.3 percent '
        'compared to the same period last year, driven primarily by strong '
        'performance in the enterprise software segment and expanding '
        'international partnerships.'
    )

    doc.add_paragraph(
        'Our customer acquisition strategy yielded significant results, with '
        '2,847 new enterprise clients onboarded during the quarter. The average '
        'deal size grew to $127,500, representing a 12 percent increase from '
        'the previous quarter. Key wins include contracts with Meridian Healthcare '
        'Systems, Vanguard Manufacturing Group, and Pacific Rim Logistics.'
    )

    doc.add_paragraph(
        'Employee satisfaction scores reached an all-time high of 4.6 out of '
        '5.0, reflecting the success of our workplace culture initiatives. The '
        'engineering team completed 94 percent of planned sprint deliverables, '
        'while the sales organization exceeded its quarterly target by 7.2 percent.'
    )

    doc.add_paragraph(
        'Looking ahead, the executive team has approved a capital expenditure '
        'budget of $14.2 million for infrastructure modernization. This investment '
        'will support the deployment of next-generation cloud services and the '
        'expansion of our data center capacity in the Asia-Pacific region.'
    )

    doc.add_paragraph(
        'Risk management assessments indicate stable market conditions with '
        'moderate exposure to currency fluctuations in European markets. The '
        'compliance team has completed the annual audit with zero material '
        'findings, maintaining our excellent regulatory standing.'
    )

    # =============================================
    # APPENDIX A (German text)
    # =============================================
    doc.add_heading('Appendix A', level=1)

    doc.add_paragraph(
        'Die Geschäftsentwicklung im dritten Quartal zeigt eine deutliche '
        'Verbesserung gegenüber dem Vorjahreszeitraum. Der Gesamtumsatz stieg '
        'um 18,3 Prozent auf 42,7 Millionen Euro. Besonders erfreulich ist das '
        'Wachstum im Bereich der Unternehmenssoftware, das maßgeblich durch '
        'neue Partnerschaften in der DACH-Region vorangetrieben wurde.'
    )

    doc.add_paragraph(
        'Die Kundenzufriedenheit erreichte einen neuen Höchstwert von 4,6 von '
        '5,0 Punkten. Unsere Mitarbeiterinnen und Mitarbeiter haben hervorragende '
        'Arbeit geleistet, insbesondere das Entwicklungsteam, das 94 Prozent '
        'aller geplanten Meilensteine termingerecht erreicht hat. Die '
        'Vertriebsabteilung übertraf ihr Quartalsziel um 7,2 Prozent.'
    )

    doc.add_paragraph(
        'Für das kommende Quartal plant die Geschäftsleitung Investitionen in '
        'Höhe von 14,2 Millionen Euro für die Modernisierung der technischen '
        'Infrastruktur. Diese Mittel werden für den Ausbau unserer '
        'Cloud-Dienste und die Erweiterung der Rechenzentrumskapazitäten im '
        'asiatisch-pazifischen Raum eingesetzt.'
    )

    # =============================================
    # APPENDIX B (Spanish text)
    # =============================================
    doc.add_heading('Appendix B', level=1)

    doc.add_paragraph(
        'El rendimiento empresarial del tercer trimestre muestra una mejora '
        'significativa en comparación con el mismo período del año anterior. '
        'Los ingresos totales aumentaron un 18,3 por ciento hasta alcanzar '
        'los 42,7 millones de euros. El crecimiento en el segmento de software '
        'empresarial fue especialmente notable, impulsado por nuevas alianzas '
        'estratégicas en la región de América Latina.'
    )

    doc.add_paragraph(
        'La satisfacción del cliente alcanzó un máximo histórico de 4,6 sobre '
        '5,0 puntos. Nuestros empleados han realizado un trabajo excepcional, '
        'particularmente el equipo de desarrollo, que completó el 94 por ciento '
        'de todos los hitos planificados dentro del plazo establecido. El '
        'departamento de ventas superó su objetivo trimestral en un 7,2 por ciento.'
    )

    doc.add_paragraph(
        'Para el próximo trimestre, la dirección ejecutiva ha aprobado una '
        'inversión de 14,2 millones de euros destinada a la modernización de '
        'la infraestructura tecnológica. Estos fondos se utilizarán para ampliar '
        'nuestros servicios en la nube y expandir la capacidad de los centros '
        'de datos en la región de Asia-Pacífico.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
