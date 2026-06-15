"""
Initial Setup: Study guide rules document with 6 rules in one continuous paragraph
Task ID: osworld_writer_spacing_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_spacing_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Study Guide: Academic Writing Rules', level=1)

    # --- First paragraph: 6 rules squashed together in ONE paragraph ---
    # IMPORTANT: All 6 rules are in a single paragraph block (no empty lines between them)
    first_para = doc.add_paragraph(
        'Always cite your sources. '
        'Use double spacing throughout. '
        'Margins must be 1 inch on all sides. '
        'Include a title page. '
        'Number all pages. '
        'Submit via the course portal.'
    )
    first_para.paragraph_format.space_before = Pt(0)
    first_para.paragraph_format.space_after = Pt(0)

    # --- Section 2: Formatting Guidelines ---
    doc.add_heading('Formatting Guidelines', level=2)

    guidelines_intro = doc.add_paragraph(
        'Proper formatting ensures your work meets academic standards and is easy for your professor to grade. '
        'The following guidelines apply to all written assignments submitted in this course.'
    )

    guidelines_detail = doc.add_paragraph(
        'Font size should be 12pt Times New Roman or Calibri. '
        'All text must be left-aligned unless otherwise instructed. '
        'Headers and subheadings should be clearly distinguished from body text using bold or a larger font size.'
    )

    # --- Section 3: Citation Requirements ---
    doc.add_heading('Citation Requirements', level=2)

    citation_intro = doc.add_paragraph(
        'Academic integrity is a cornerstone of scholarly work at this institution. '
        'All sources referenced in your paper must be properly cited according to the style guide specified in your syllabus.'
    )

    citation_types = doc.add_paragraph(
        'In-text citations must include the author\'s last name and year of publication. '
        'Direct quotations require page numbers in addition to the author and year. '
        'Paraphrased material still requires a citation even though it is not a direct quote.'
    )

    # --- Section 4: Submission Instructions ---
    doc.add_heading('Submission Instructions', level=2)

    submission_para = doc.add_paragraph(
        'All assignments must be submitted electronically through the course portal before the deadline. '
        'Late submissions will incur a 10% penalty per day unless a prior extension has been granted in writing. '
        'Technical issues are not an acceptable excuse for late submission, so plan accordingly and submit early.'
    )

    submission_format = doc.add_paragraph(
        'Save your file as a .docx or .pdf document with your last name and assignment title in the filename. '
        'For example: Smith_ResearchPaper.docx. '
        'Ensure your student ID number appears in the header of every page.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
