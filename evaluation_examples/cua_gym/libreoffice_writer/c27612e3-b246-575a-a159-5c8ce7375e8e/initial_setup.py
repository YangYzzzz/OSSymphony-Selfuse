"""
Initial Setup: Create a legal contract document with 12 sections.
Task ID: writer_legal_094
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_094'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_section_heading(doc, text):
    """Add a bold section heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    return para


def add_body_text(doc, text):
    """Add body paragraph with standard formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return para


def create_initial():
    doc = Document()

    # Set default margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run('MASTER SERVICES AGREEMENT')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    # Subtitle / date line
    sub = doc.add_paragraph()
    sub.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.paragraph_format.space_after = Pt(12)
    run = sub.add_run('Effective Date: March 15, 2025')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    # Parties preamble
    add_body_text(doc,
        'This Master Services Agreement ("Agreement") is entered into by and between '
        'Meridian Technology Solutions, Inc., a Delaware corporation with its principal '
        'office at 2400 Embarcadero Way, Palo Alto, CA 94303 ("Provider"), and '
        'Cornerstone Financial Group, LLC, a New York limited liability company with '
        'its principal office at 780 Third Avenue, Suite 4200, New York, NY 10017 ("Client").'
    )

    add_body_text(doc,
        'WHEREAS, Client desires to engage Provider to perform certain professional '
        'services as described herein, and Provider desires to perform such services, '
        'subject to the terms and conditions set forth in this Agreement.'
    )

    # --- Section 1: Definitions ---
    add_section_heading(doc, 'Section 1 - Definitions')
    add_body_text(doc,
        '"Confidential Information" means any non-public information disclosed by '
        'either party to the other, whether orally, in writing, or by inspection of '
        'tangible objects, that is designated as confidential or that reasonably should '
        'be understood to be confidential given the nature of the information and the '
        'circumstances of disclosure. See Section 5 for confidentiality obligations.'
    )
    add_body_text(doc,
        '"Deliverables" means all work product, documentation, software, reports, and '
        'other materials produced by Provider in the course of performing the Services, '
        'as further described in the applicable Statement of Work. Acceptance of '
        'Deliverables is governed by Section 3.'
    )
    add_body_text(doc,
        '"Services" means the professional services to be performed by Provider as '
        'described in Section 2 and any applicable Statement of Work executed by both parties.'
    )

    # --- Section 2: Scope of Services ---
    add_section_heading(doc, 'Section 2 - Scope of Services')
    add_body_text(doc,
        'Provider shall perform the services described in each Statement of Work '
        '("SOW") executed by both parties. Each SOW shall specify the scope, timeline, '
        'milestones, deliverables, and fees for the applicable engagement. Provider '
        'shall assign qualified personnel to perform the Services and shall maintain '
        'adequate staffing levels throughout the engagement period.'
    )
    add_body_text(doc,
        'Provider shall perform all Services in a professional and workmanlike manner, '
        'consistent with industry standards. Provider shall comply with all applicable '
        'laws and regulations in the performance of the Services. Any failure to meet '
        'the service standards described herein shall be subject to the remedies set '
        'forth in Section 7 and Section 10.'
    )

    # --- Section 3: Deliverables and Acceptance ---
    add_section_heading(doc, 'Section 3 - Deliverables and Acceptance')
    add_body_text(doc,
        'Upon completion of each Deliverable, Provider shall submit the Deliverable '
        'to Client for review and acceptance. Client shall have fifteen (15) business '
        'days following receipt to review each Deliverable and either accept or reject '
        'it in writing. If Client rejects a Deliverable, Client shall provide written '
        'notice specifying in reasonable detail the deficiencies that must be corrected.'
    )
    add_body_text(doc,
        'Provider shall use commercially reasonable efforts to correct any deficiencies '
        'within ten (10) business days of receiving such notice. If a dispute arises '
        'regarding acceptance, the parties shall follow the dispute resolution process '
        'described in Section 11.'
    )

    # --- Section 4: Compensation and Payment ---
    add_section_heading(doc, 'Section 4 - Compensation and Payment')
    add_body_text(doc,
        'Client shall pay Provider the fees specified in each SOW. Unless otherwise '
        'stated in the applicable SOW, Provider shall invoice Client monthly for '
        'Services performed during the preceding calendar month. Payment shall be '
        'due within thirty (30) days of receipt of a proper invoice. Late payments '
        'shall accrue interest at the rate of 1.5% per month or the maximum rate '
        'permitted by applicable law, whichever is less.'
    )
    add_body_text(doc,
        'Client shall reimburse Provider for reasonable, pre-approved out-of-pocket '
        'expenses incurred in the performance of the Services, provided that Provider '
        'submits appropriate documentation. The total fees payable under this Agreement '
        'shall not exceed the aggregate cap specified in the applicable SOW without '
        'prior written approval from Client.'
    )

    # --- Section 5: Confidentiality ---
    add_section_heading(doc, 'Section 5 - Confidentiality')
    add_body_text(doc,
        'Each party agrees to hold the other party\'s Confidential Information in '
        'strict confidence and not to disclose it to any third party without the prior '
        'written consent of the disclosing party. Each party shall use at least the '
        'same degree of care to protect the other party\'s Confidential Information as '
        'it uses to protect its own confidential information, but in no event less than '
        'reasonable care. This obligation survives termination for a period of three (3) '
        'years.'
    )
    add_body_text(doc,
        'Confidential Information shall not include information that: (a) is or becomes '
        'publicly available through no fault of the receiving party; (b) was known to '
        'the receiving party prior to disclosure, as evidenced by written records; '
        '(c) is independently developed by the receiving party without use of or '
        'reference to the disclosing party\'s Confidential Information; or (d) is '
        'rightfully obtained from a third party without restriction on disclosure. '
        'Any breach of this Section 5 may result in remedies under Section 8 in '
        'addition to any other remedies available at law or in equity.'
    )

    # --- Section 6: Intellectual Property ---
    add_section_heading(doc, 'Section 6 - Intellectual Property')
    add_body_text(doc,
        'All Deliverables created by Provider in the performance of the Services shall '
        'be considered works made for hire and shall be the exclusive property of Client. '
        'To the extent that any Deliverable does not qualify as a work made for hire, '
        'Provider hereby assigns to Client all right, title, and interest in and to '
        'such Deliverable, including all intellectual property rights therein.'
    )
    add_body_text(doc,
        'Provider retains ownership of all pre-existing intellectual property, tools, '
        'methodologies, and frameworks that Provider brings to the engagement ("Provider '
        'IP"). Provider grants Client a non-exclusive, perpetual, royalty-free license '
        'to use any Provider IP incorporated into the Deliverables solely to the extent '
        'necessary for Client to use the Deliverables for their intended purpose.'
    )

    # --- Section 7: Limitation of Liability ---
    add_section_heading(doc, 'Section 7 - Limitation of Liability')
    add_body_text(doc,
        'EXCEPT FOR BREACHES OF SECTION 5 (CONFIDENTIALITY) OR SECTION 8 '
        '(INDEMNIFICATION), NEITHER PARTY SHALL BE LIABLE TO THE OTHER FOR ANY '
        'INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING '
        'OUT OF OR RELATED TO THIS AGREEMENT, REGARDLESS OF WHETHER SUCH DAMAGES '
        'ARE BASED ON CONTRACT, TORT, STRICT LIABILITY, OR ANY OTHER THEORY, EVEN '
        'IF THE PARTY HAS BEEN ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.'
    )
    add_body_text(doc,
        'EXCEPT FOR BREACHES OF SECTION 5 (CONFIDENTIALITY) OR OBLIGATIONS UNDER '
        'SECTION 8 (INDEMNIFICATION), EACH PARTY\'S TOTAL AGGREGATE LIABILITY UNDER '
        'THIS AGREEMENT SHALL NOT EXCEED THE TOTAL FEES PAID OR PAYABLE BY CLIENT '
        'TO PROVIDER UNDER THE APPLICABLE SOW DURING THE TWELVE (12) MONTH PERIOD '
        'PRECEDING THE EVENT GIVING RISE TO THE CLAIM.'
    )
    add_body_text(doc,
        'The limitations set forth in this Section 7 shall apply to the fullest extent '
        'permitted by applicable law and shall survive the termination or expiration of '
        'this Agreement. Nothing in this Section 7 shall limit either party\'s liability '
        'for fraud, gross negligence, willful misconduct, or death or personal injury '
        'caused by negligence.'
    )

    # --- Section 8: Indemnification ---
    add_section_heading(doc, 'Section 8 - Indemnification')
    add_body_text(doc,
        'Provider shall indemnify, defend, and hold harmless Client and its officers, '
        'directors, employees, agents, and affiliates from and against any and all '
        'third-party claims, damages, losses, liabilities, costs, and expenses '
        '(including reasonable attorneys\' fees) arising out of or related to: '
        '(a) Provider\'s breach of any representation, warranty, or obligation under '
        'this Agreement; (b) any negligent or wrongful act or omission of Provider '
        'or its personnel; or (c) any claim that the Deliverables infringe or '
        'misappropriate any third party\'s intellectual property rights.'
    )
    add_body_text(doc,
        'Client shall indemnify, defend, and hold harmless Provider and its officers, '
        'directors, employees, agents, and affiliates from and against any and all '
        'third-party claims, damages, losses, liabilities, costs, and expenses '
        '(including reasonable attorneys\' fees) arising out of or related to: '
        '(a) Client\'s breach of any representation, warranty, or obligation under '
        'this Agreement; or (b) Client\'s use of the Deliverables in a manner not '
        'authorized by this Agreement or the applicable SOW.'
    )
    add_body_text(doc,
        'The indemnifying party\'s obligations under this Section 8 are conditioned '
        'upon the indemnified party: (i) promptly notifying the indemnifying party '
        'in writing of any claim; (ii) granting the indemnifying party sole control '
        'of the defense and settlement of the claim; and (iii) providing reasonable '
        'assistance and cooperation at the indemnifying party\'s expense. The '
        'indemnification obligations in this Section 8 are subject to the limitations '
        'set forth in Section 7, except that the aggregate liability cap shall not '
        'apply to indemnification claims arising from intellectual property infringement.'
    )

    # --- Section 9: Term and Termination ---
    add_section_heading(doc, 'Section 9 - Term and Termination')
    add_body_text(doc,
        'This Agreement shall commence on the Effective Date and shall continue for '
        'an initial term of three (3) years, unless earlier terminated in accordance '
        'with this Section 9. The Agreement shall automatically renew for successive '
        'one (1) year periods unless either party provides written notice of non-renewal '
        'at least ninety (90) days prior to the end of the then-current term.'
    )
    add_body_text(doc,
        'Either party may terminate this Agreement or any SOW for cause upon thirty '
        '(30) days\' written notice if the other party materially breaches this '
        'Agreement and fails to cure such breach within the notice period. Client '
        'may terminate any SOW for convenience upon sixty (60) days\' written notice, '
        'subject to payment for all Services performed through the effective date of '
        'termination. Upon termination, the provisions of Section 5, Section 6, '
        'Section 7, Section 8, Section 11, and Section 12 shall survive.'
    )

    # --- Section 10: Representations and Warranties ---
    add_section_heading(doc, 'Section 10 - Representations and Warranties')
    add_body_text(doc,
        'Provider represents and warrants that: (a) it has the legal right and '
        'authority to enter into this Agreement and perform the Services; (b) the '
        'Services will be performed in a professional manner consistent with '
        'generally accepted industry standards; (c) the Deliverables will conform '
        'in all material respects to the specifications set forth in the applicable '
        'SOW; and (d) to the best of Provider\'s knowledge, the Deliverables will '
        'not infringe any third party\'s intellectual property rights.'
    )
    add_body_text(doc,
        'EXCEPT AS EXPRESSLY SET FORTH IN THIS AGREEMENT, PROVIDER MAKES NO '
        'WARRANTIES, EXPRESS OR IMPLIED, INCLUDING WITHOUT LIMITATION ANY IMPLIED '
        'WARRANTIES OF MERCHANTABILITY OR FITNESS FOR A PARTICULAR PURPOSE. Client\'s '
        'exclusive remedy for breach of the warranties in this Section 10 shall be, '
        'at Provider\'s option, re-performance of the applicable Services or a refund '
        'of the fees paid for the non-conforming Deliverable, subject to the limitations '
        'in Section 7.'
    )

    # --- Section 11: Dispute Resolution ---
    add_section_heading(doc, 'Section 11 - Dispute Resolution')
    add_body_text(doc,
        'Any dispute, controversy, or claim arising out of or relating to this '
        'Agreement shall first be submitted to good faith negotiation between senior '
        'executives of each party. If the dispute is not resolved within thirty (30) '
        'days of written notice, either party may submit the dispute to binding '
        'arbitration administered by the American Arbitration Association under its '
        'Commercial Arbitration Rules. The arbitration shall be conducted in New York, '
        'New York, and the decision of the arbitrator shall be final and binding.'
    )
    add_body_text(doc,
        'Notwithstanding the foregoing, either party may seek injunctive or other '
        'equitable relief in any court of competent jurisdiction to prevent irreparable '
        'harm pending resolution of the dispute, including but not limited to breaches '
        'of Section 5 (Confidentiality) or Section 6 (Intellectual Property). The '
        'prevailing party in any arbitration or litigation shall be entitled to recover '
        'its reasonable attorneys\' fees and costs.'
    )

    # --- Section 12: General Provisions ---
    add_section_heading(doc, 'Section 12 - General Provisions')
    add_body_text(doc,
        'This Agreement constitutes the entire agreement between the parties with '
        'respect to the subject matter hereof and supersedes all prior and '
        'contemporaneous agreements, proposals, negotiations, representations, and '
        'communications, whether oral or written. This Agreement may only be amended '
        'by a written instrument executed by authorized representatives of both parties.'
    )
    add_body_text(doc,
        'Neither party may assign this Agreement or any rights or obligations hereunder '
        'without the prior written consent of the other party, except that either party '
        'may assign this Agreement to an affiliate or in connection with a merger, '
        'acquisition, or sale of all or substantially all of its assets. Any attempted '
        'assignment in violation of this Section 12 shall be void.'
    )
    add_body_text(doc,
        'If any provision of this Agreement is held to be invalid or unenforceable, '
        'the remaining provisions shall continue in full force and effect. The failure '
        'of either party to enforce any right or provision of this Agreement shall not '
        'constitute a waiver of such right or provision. This Agreement shall be '
        'governed by and construed in accordance with the laws of the State of New York, '
        'without regard to its conflict of laws principles. All notices under this '
        'Agreement shall be in writing and delivered to the addresses set forth in '
        'Section 12 or such other address as a party may designate in writing.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
