#!/usr/bin/env python3
"""
initial_setup.py - Create a 5-page FAQ document with 20 Q&A pairs.
All paragraphs have default text flow settings (no keep_together).
"""
import os
import math
import subprocess
import shlex
import time

# Install dependency
os.system("pip3 install python-docx 2>/dev/null")

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

CHARS_PER_LINE = 85  # proxy for rendered line width

def estimate_lines(text):
    return math.ceil(len(text) / CHARS_PER_LINE)

# 20 Q&A pairs with realistic tech/business FAQ content
# Carefully crafted so paragraph lengths range 1-8 lines
# Some < 4 lines (short), some >= 4 lines (long)
qa_pairs = [
    # Q1 - short answer (~1 line, ~40 chars)
    ("What are your business hours?",
     "Our offices are open Monday through Friday, 9 AM to 6 PM Eastern Time."),

    # Q2 - short answer (~2 lines, ~130 chars)
    ("How do I reset my password?",
     "To reset your password, click the 'Forgot Password' link on the login page. You will receive an email with a secure link to create a new password within 15 minutes."),

    # Q3 - long answer (~4 lines, ~340 chars)
    ("What payment methods do you accept?",
     "We accept all major credit cards including Visa, MasterCard, American Express, and Discover. We also support payment through PayPal, Apple Pay, and Google Pay for online transactions. For enterprise customers, we offer invoicing with net-30 payment terms. Bank wire transfers are available for orders exceeding $5,000. Please contact our billing department for custom payment arrangements."),

    # Q4 - short answer (~1 line, ~60 chars)
    ("Is there a free trial available?",
     "Yes, we offer a 14-day free trial with full access to all features."),

    # Q5 - long answer (~5 lines, ~420 chars)
    ("What is your refund policy?",
     "We offer a 30-day money-back guarantee on all subscription plans. If you are not satisfied with our service, you can request a full refund within the first 30 days of your purchase. After the initial 30-day period, refunds are prorated based on the remaining time in your billing cycle. To initiate a refund, please contact our support team with your order number and reason for cancellation. Enterprise contracts may have different refund terms as specified in the agreement."),

    # Q6 - short answer (~2 lines, ~150 chars)
    ("How do I contact customer support?",
     "You can reach our customer support team by email at support@example.com, by phone at 1-800-555-0123, or through the live chat widget on our website during business hours."),

    # Q7 - long answer (~6 lines, ~510 chars)
    ("How does data encryption work in your platform?",
     "All data transmitted between your browser and our servers is encrypted using TLS 1.3 with 256-bit AES encryption. Data at rest is encrypted using AES-256-GCM with keys managed through our hardware security modules. We implement envelope encryption where data encryption keys are themselves encrypted by master keys stored in tamper-resistant hardware. Additionally, all database backups are encrypted before being stored in geographically distributed locations. Our encryption practices are regularly audited and comply with SOC 2 Type II and ISO 27001 standards."),

    # Q8 - short answer (~3 lines, ~230 chars)
    ("Can I upgrade my plan at any time?",
     "Yes, you can upgrade your subscription plan at any time from your account settings dashboard. When you upgrade, you will be charged a prorated amount for the remainder of your current billing period. The new features become available immediately after upgrading."),

    # Q9 - long answer (~5 lines, ~400 chars)
    ("What integrations are available?",
     "Our platform integrates with over 200 third-party applications including Salesforce, HubSpot, Slack, Microsoft Teams, Jira, and Trello. We provide native REST APIs and webhooks for custom integrations. Our marketplace also features pre-built connectors for popular CRM, ERP, and project management tools. For enterprise customers, we offer dedicated integration support and custom connector development to ensure seamless workflow automation across your technology stack."),

    # Q10 - short answer (~1 line, ~55 chars)
    ("Do you offer multi-language support?",
     "Yes, our platform is available in 24 languages worldwide."),

    # Q11 - long answer (~7 lines, ~580 chars)
    ("What are the system requirements for the desktop application?",
     "The desktop application requires Windows 10 or later, macOS 11 Big Sur or later, or Ubuntu 20.04 LTS or later. A minimum of 4 GB of RAM is recommended, though 8 GB provides optimal performance for large datasets. You will need at least 500 MB of free disk space for installation and an additional 2 GB for local data caching. A stable internet connection with at least 5 Mbps download speed is required for real-time synchronization features. The application supports both Intel and ARM-based processors on all platforms. Graphics hardware acceleration is optional but recommended for the dashboard visualization components."),

    # Q12 - short answer (~2 lines, ~160 chars)
    ("How often do you release updates?",
     "We release minor updates and bug fixes on a bi-weekly basis. Major feature releases occur quarterly, and all updates are applied automatically unless you opt for manual update management."),

    # Q13 - long answer (~4 lines, ~330 chars)
    ("What compliance certifications do you hold?",
     "We maintain SOC 2 Type II, ISO 27001, and GDPR compliance certifications. Our platform is also HIPAA-compliant for healthcare customers and PCI DSS Level 1 certified for payment processing. We undergo annual third-party security audits and penetration testing. Compliance reports and certificates are available upon request through your account manager."),

    # Q14 - short answer (~3 lines, ~240 chars)
    ("Can I export my data?",
     "You can export all of your data at any time in multiple formats including CSV, JSON, XML, and PDF. Bulk export options are available from the administration panel. We also provide a Data Portability API for automated data extraction and migration to other platforms."),

    # Q15 - long answer (~8 lines, ~650 chars)
    ("How does your disaster recovery process work?",
     "Our disaster recovery infrastructure ensures maximum uptime and data protection through multiple redundant systems. We maintain real-time database replication across three geographically separated data centers, with automatic failover capabilities that activate within 60 seconds of a primary system failure. Full system backups are performed every four hours and retained for 90 days, while incremental backups run every 15 minutes. Our Recovery Point Objective is less than 15 minutes, and our Recovery Time Objective is under 4 hours for a complete regional failure scenario. We conduct quarterly disaster recovery drills to validate our procedures and update our runbooks accordingly. All recovery processes are documented and certified as part of our SOC 2 compliance program."),

    # Q16 - short answer (~1 line, ~70 chars)
    ("Is there a mobile app available?",
     "Yes, native mobile apps are available for both iOS and Android on their app stores."),

    # Q17 - long answer (~5 lines, ~430 chars)
    ("What training resources do you provide?",
     "We offer a comprehensive learning center with video tutorials, step-by-step guides, and interactive walkthroughs covering all platform features. New users receive a guided onboarding experience with personalized setup assistance during their first week. Weekly live webinars cover advanced topics and best practices. Enterprise customers receive dedicated training sessions and can schedule custom workshops for their teams. Our community forum also provides peer-to-peer support and knowledge sharing opportunities."),

    # Q18 - short answer (~2 lines, ~140 chars)
    ("What happens when my trial expires?",
     "When your trial expires, your account will be converted to a limited free plan. All your data is preserved, and you can upgrade to a paid plan at any time."),

    # Q19 - long answer (~6 lines, ~500 chars)
    ("How do you handle service outages?",
     "We maintain a public status page at status.example.com that provides real-time updates on all system components. When an incident occurs, our on-call engineering team is automatically alerted and begins investigation within five minutes. We classify incidents by severity level and communicate updates every 30 minutes through the status page, email notifications, and in-app alerts. Post-incident reviews are conducted within 48 hours and published as public postmortems. Our historical uptime exceeds 99.95 percent across all services."),

    # Q20 - short answer (~3 lines, ~220 chars)
    ("Do you offer volume discounts?",
     "Yes, we offer tiered volume discounts for teams of 10 or more users. Discounts range from 10 to 30 percent depending on team size and contract length. Contact our sales team for a customized enterprise quote."),
]

# Verify line estimates
for i, (q, a) in enumerate(qa_pairs):
    lines = estimate_lines(a)
    print(f"Q{i+1}: '{q[:40]}...' -> {len(a)} chars, ~{lines} lines")

# Build document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Set margins for standard page width
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Title
title = doc.add_heading('Frequently Asked Questions', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run('Product Documentation & Support Guide')
run.font.size = Pt(13)
run.font.color.rgb = None  # default color
run.italic = True

doc.add_paragraph()  # spacing

# Add Q&A pairs
for i, (question, answer) in enumerate(qa_pairs):
    # Question as heading
    doc.add_heading(f'Q{i+1}. {question}', level=2)

    # Answer as body paragraph - default text flow (no keep_together set)
    para = doc.add_paragraph(answer)
    para.paragraph_format.space_after = Pt(6)

# Save
output_path = '/home/user/wrpara_036.docx'
doc.save(output_path)
print(f"\nDocument saved to {output_path}")

# Verify no keep_together is set
doc2 = Document(output_path)
for i, para in enumerate(doc2.paragraphs):
    kt = para.paragraph_format.keep_together
    if kt is not None and kt is not False:
        print(f"WARNING: Paragraph {i} has keep_together={kt}")

print("Verification complete - all paragraphs have default text flow settings.")

# Launch LibreOffice Writer
env = os.environ.copy()
env["DISPLAY"] = ":0"
subprocess.Popen(
    shlex.split(f'libreoffice --writer "{output_path}"'),
    stdout=subprocess.DEVNULL,
    stderr=subprocess.DEVNULL,
    env=env,
)
time.sleep(2)
print("LibreOffice Writer launched.")
