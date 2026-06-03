"""
Reward Script: Dissertation front matter — ToC, List of Figures, List of Tables, List of Abbreviations
Task ID: writer_mt_095
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Table of Contents section exists with Heading 1 title and entries
  Component 2 (0.25): List of Figures section exists with Heading 1 title and 15 figure entries
  Component 3 (0.20): List of Tables section exists with Heading 1 title and 10 table entries
  Component 4 (0.15): List of Abbreviations section exists with Heading 1 title and abbreviation entries
  Component 5 (0.15): Page breaks separate each front matter section (at least 3 new page breaks)
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_095'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def parse_front_matter(doc):
    """
    Parse the document to identify front matter sections inserted between
    the title page and the first chapter heading.
    Returns a dict with keys: 'toc', 'figures', 'tables', 'abbreviations'.
    Each value is a list of entry texts found in that section.
    Also returns the heading paragraph indices for each section.
    """
    sections = {
        'toc': {'title_found': False, 'entries': [], 'heading_idx': None},
        'figures': {'title_found': False, 'entries': [], 'heading_idx': None},
        'tables': {'title_found': False, 'entries': [], 'heading_idx': None},
        'abbreviations': {'title_found': False, 'entries': [], 'heading_idx': None},
    }

    current_section = None
    title_map = {
        'Table of Contents': 'toc',
        'List of Figures': 'figures',
        'List of Tables': 'tables',
        'List of Abbreviations': 'abbreviations',
    }

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style_name = p.style.name

        # Detect front matter Heading 1 titles
        if style_name == 'Heading 1':
            if text in title_map:
                key = title_map[text]
                sections[key]['title_found'] = True
                sections[key]['heading_idx'] = i
                current_section = key
                continue
            elif text.startswith('Chapter') or text == 'Glossary of Terms':
                # We've exited front matter into the body
                current_section = None
                break
            else:
                current_section = None
                continue

        # Collect entries under current front matter section
        if current_section and text:
            sections[current_section]['entries'].append(text)

    return sections


def count_page_breaks_in_range(doc, start_idx, end_idx):
    """Count inline page breaks in paragraphs between start_idx and end_idx (exclusive)."""
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    count = 0
    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        # Check page_break_before property
        if p.paragraph_format.page_break_before:
            count += 1
        # Check inline page breaks in runs
        for run in p.runs:
            for br in run.element.findall(".//{%s}br" % ns_w):
                bt = br.attrib.get("{%s}type" % ns_w, "")
                if bt == "page":
                    count += 1
    return count


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    fm = parse_front_matter(doc)

    # Component 1: Table of Contents (0.25 points)
    # Must have Heading 1 "Table of Contents" with TOC entries listing chapters/sections
    try:
        if fm['toc']['title_found']:
            toc_entries = fm['toc']['entries']
            # TOC should have chapter entries (at least 8 chapters + some subsections)
            chapter_entries = [e for e in toc_entries if 'Chapter' in e or any(
                e.startswith(f'{n}.') for n in range(1, 10)
            )]
            if len(chapter_entries) >= 8:
                print(f"PASS: Component 1 — Table of Contents found with {len(toc_entries)} entries "
                      f"({len(chapter_entries)} chapter/section entries) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — TOC title found but only {len(chapter_entries)} "
                      f"chapter/section entries (expected >= 8)")
        else:
            print("FAIL: Component 1 — No 'Table of Contents' Heading 1 found in front matter")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: List of Figures (0.25 points)
    # Must have Heading 1 "List of Figures" with 15 figure entries
    try:
        if fm['figures']['title_found']:
            fig_entries = fm['figures']['entries']
            figure_items = [e for e in fig_entries if e.lower().startswith('figure')]
            if len(figure_items) >= 15:
                print(f"PASS: Component 2 — List of Figures found with {len(figure_items)} "
                      f"figure entries (0.25 pts)")
                total_score += 0.25
            elif len(figure_items) >= 10:
                partial = round(0.25 * len(figure_items) / 15, 2)
                print(f"PARTIAL: Component 2 — List of Figures found but only {len(figure_items)} "
                      f"figure entries (expected 15) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — List of Figures title found but only {len(figure_items)} "
                      f"figure entries (expected 15)")
        else:
            print("FAIL: Component 2 — No 'List of Figures' Heading 1 found in front matter")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: List of Tables (0.20 points)
    # Must have Heading 1 "List of Tables" with 10 table entries
    try:
        if fm['tables']['title_found']:
            tbl_entries = fm['tables']['entries']
            table_items = [e for e in tbl_entries if e.lower().startswith('table')]
            if len(table_items) >= 10:
                print(f"PASS: Component 3 — List of Tables found with {len(table_items)} "
                      f"table entries (0.20 pts)")
                total_score += 0.20
            elif len(table_items) >= 5:
                partial = round(0.20 * len(table_items) / 10, 2)
                print(f"PARTIAL: Component 3 — List of Tables found but only {len(table_items)} "
                      f"table entries (expected 10) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — List of Tables title found but only {len(table_items)} "
                      f"table entries (expected 10)")
        else:
            print("FAIL: Component 3 — No 'List of Tables' Heading 1 found in front matter")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: List of Abbreviations (0.15 points)
    # Must have Heading 1 "List of Abbreviations" with abbreviation entries
    try:
        if fm['abbreviations']['title_found']:
            abbrev_entries = fm['abbreviations']['entries']
            if len(abbrev_entries) >= 15:
                print(f"PASS: Component 4 — List of Abbreviations found with {len(abbrev_entries)} "
                      f"entries (0.15 pts)")
                total_score += 0.15
            elif len(abbrev_entries) >= 5:
                partial = round(0.15 * len(abbrev_entries) / 20, 2)
                print(f"PARTIAL: Component 4 — List of Abbreviations found but only {len(abbrev_entries)} "
                      f"entries (expected ~20) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — List of Abbreviations title found but only "
                      f"{len(abbrev_entries)} entries (expected ~20)")
        else:
            print("FAIL: Component 4 — No 'List of Abbreviations' Heading 1 found in front matter")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Page breaks between front matter sections (0.15 points)
    # The task requires each section to start on a new page.
    # Gate: at least one front matter section must exist (otherwise no front matter was added)
    try:
        fm_sections_found = sum(1 for k in ['toc', 'figures', 'tables', 'abbreviations']
                                if fm[k]['title_found'])
        if fm_sections_found == 0:
            print("FAIL: Component 5 — No front matter sections found, so page break check is moot")
        else:
            # Find the range: from just after title to first chapter heading
            first_chapter_idx = None
            for i, p in enumerate(doc.paragraphs):
                if p.style.name == 'Heading 1' and p.text.strip().startswith('Chapter'):
                    first_chapter_idx = i
                    break

            # Find title page end (the Title style paragraph is around idx 6)
            title_end_idx = 0
            for i, p in enumerate(doc.paragraphs):
                if p.style.name == 'Title':
                    title_end_idx = i + 1
                    break

            if first_chapter_idx is None:
                first_chapter_idx = len(doc.paragraphs)

            # Count page breaks in the front matter region
            ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            pb_count = 0
            for i in range(title_end_idx, first_chapter_idx):
                p = doc.paragraphs[i]
                if p.paragraph_format.page_break_before:
                    pb_count += 1
                for run in p.runs:
                    for br in run.element.findall(".//{%s}br" % ns_w):
                        bt = br.attrib.get("{%s}type" % ns_w, "")
                        if bt == "page":
                            pb_count += 1

            # We expect at least 3 page breaks to separate 4 front matter sections
            # The golden has breaks at paragraphs 10, 52, 70, 83, 106 (5 total).
            if pb_count >= 3:
                print(f"PASS: Component 5 — Found {pb_count} page breaks in front matter region "
                      f"(paragraphs {title_end_idx}-{first_chapter_idx}) (0.15 pts)")
                total_score += 0.15
            elif pb_count >= 1:
                partial = round(0.15 * pb_count / 3, 2)
                print(f"PARTIAL: Component 5 — Found {pb_count} page breaks in front matter region "
                      f"(expected >= 3) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — No page breaks found in front matter region "
                      f"(paragraphs {title_end_idx}-{first_chapter_idx})")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification (best-effort)
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
