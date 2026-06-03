"""
Initial Setup: Create a dissertation document with chapters, figures, tables,
and abbreviation entries, but NO front matter indexes.
Task ID: writer_mt_095
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from PIL import Image
import io

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_095'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_placeholder_image(path, width=400, height=250, label="Figure"):
    """Create a simple placeholder image for figures."""
    img = Image.new('RGB', (width, height), color=(220, 230, 240))
    # Draw a simple border effect
    for x in range(width):
        for y in [0, 1, height - 1, height - 2]:
            img.putpixel((x, y), (100, 120, 150))
    for y in range(height):
        for x in [0, 1, width - 1, width - 2]:
            img.putpixel((x, y), (100, 120, 150))
    img.save(path)


# Chapter data: 8 chapters with sub-sections distributed among them
CHAPTERS = [
    {
        "title": "Introduction",
        "subsections": [
            "Background and Motivation",
            "Research Questions",
            "Scope and Limitations",
            "Thesis Structure",
        ],
        "figures": [1],
        "tables": [1],
    },
    {
        "title": "Literature Review",
        "subsections": [
            "Theoretical Foundations",
            "Previous Approaches to Neural Architecture Search",
            "Transfer Learning in Low-Resource Settings",
            "Gaps in Existing Research",
        ],
        "figures": [2, 3],
        "tables": [2],
    },
    {
        "title": "Methodology",
        "subsections": [
            "Research Design",
            "Data Collection Procedures",
            "Model Architecture",
            "Training Protocol and Hyperparameters",
            "Evaluation Metrics",
        ],
        "figures": [4, 5],
        "tables": [3, 4],
    },
    {
        "title": "Dataset Construction and Preprocessing",
        "subsections": [
            "Source Corpora",
            "Annotation Guidelines",
            "Inter-Annotator Agreement",
            "Data Augmentation Strategies",
        ],
        "figures": [6, 7],
        "tables": [5],
    },
    {
        "title": "Experimental Setup",
        "subsections": [
            "Baseline Models",
            "Hardware and Software Environment",
            "Cross-Validation Strategy",
        ],
        "figures": [8],
        "tables": [6],
    },
    {
        "title": "Results and Analysis",
        "subsections": [
            "Quantitative Results",
            "Ablation Study",
            "Error Analysis",
            "Statistical Significance Tests",
            "Comparison with State-of-the-Art",
        ],
        "figures": [9, 10, 11, 12],
        "tables": [7, 8],
    },
    {
        "title": "Discussion",
        "subsections": [
            "Interpretation of Findings",
            "Practical Implications",
            "Theoretical Contributions",
        ],
        "figures": [13, 14],
        "tables": [9],
    },
    {
        "title": "Conclusion and Future Work",
        "subsections": [
            "Summary of Contributions",
            "Future Research Directions",
        ],
        "figures": [15],
        "tables": [10],
    },
]

# Figure captions
FIGURE_CAPTIONS = [
    "Overview of the proposed research framework",
    "Timeline of related work in neural architecture search",
    "Taxonomy of transfer learning approaches",
    "High-level system architecture diagram",
    "Detailed model pipeline with preprocessing stages",
    "Distribution of samples across source corpora",
    "Annotation interface screenshot with example labels",
    "Baseline model performance on validation set",
    "Accuracy comparison across all experimental conditions",
    "Learning curves for the top three model configurations",
    "Confusion matrix for the best-performing model",
    "ROC curves with AUC scores for each class",
    "Relationship between model complexity and performance",
    "Deployment architecture for real-world application",
    "Proposed extensions for future multi-modal integration",
]

# Table captions
TABLE_CAPTIONS = [
    "Summary of research questions and methods",
    "Comparison of existing approaches in the literature",
    "Hyperparameter search space and selected values",
    "Dataset statistics after preprocessing",
    "Inter-annotator agreement scores by category",
    "Hardware and software specifications",
    "Main experimental results with confidence intervals",
    "Ablation study results removing individual components",
    "Practical deployment requirements and constraints",
    "Summary of contributions and corresponding chapters",
]

# Abbreviations (20 entries)
ABBREVIATIONS = [
    ("NAS", "Neural Architecture Search"),
    ("NLP", "Natural Language Processing"),
    ("CNN", "Convolutional Neural Network"),
    ("RNN", "Recurrent Neural Network"),
    ("LSTM", "Long Short-Term Memory"),
    ("GRU", "Gated Recurrent Unit"),
    ("BERT", "Bidirectional Encoder Representations from Transformers"),
    ("GPT", "Generative Pre-trained Transformer"),
    ("SVM", "Support Vector Machine"),
    ("ROC", "Receiver Operating Characteristic"),
    ("AUC", "Area Under the Curve"),
    ("F1", "F1 Score (Harmonic Mean of Precision and Recall)"),
    ("BLEU", "Bilingual Evaluation Understudy"),
    ("PCA", "Principal Component Analysis"),
    ("SGD", "Stochastic Gradient Descent"),
    ("API", "Application Programming Interface"),
    ("GPU", "Graphics Processing Unit"),
    ("TPU", "Tensor Processing Unit"),
    ("EDA", "Exploratory Data Analysis"),
    ("IoU", "Intersection over Union"),
]

# Body text templates for paragraphs
BODY_TEXTS = {
    "Introduction": [
        "The rapid advancement of deep learning techniques has transformed the landscape of computational linguistics and natural language processing. Over the past decade, the development of increasingly sophisticated neural architectures has enabled breakthroughs in tasks ranging from machine translation to question answering.",
        "This dissertation investigates the intersection of neural architecture search and transfer learning, with a particular focus on low-resource language settings. The primary motivation stems from the observation that manually designed architectures often underperform when adapted to typologically diverse languages.",
    ],
    "Literature Review": [
        "The theoretical foundations of this work draw upon three interconnected research streams: representation learning, meta-learning, and cross-lingual transfer. Bengio et al. (2013) established the foundational framework for learning distributed representations, which has since been extended by numerous subsequent studies.",
        "Previous approaches to neural architecture search have primarily focused on high-resource scenarios where large amounts of labeled training data are readily available. Zoph and Le (2017) demonstrated that reinforcement learning could effectively discover architectures that rival or surpass human-designed counterparts.",
    ],
    "Methodology": [
        "The research design follows a mixed-methods approach combining quantitative experimentation with qualitative error analysis. The quantitative component evaluates model performance across multiple benchmark datasets, while the qualitative analysis examines systematic patterns in model failures.",
        "The proposed model architecture consists of a hierarchical encoder with attention mechanisms at both the token and sentence levels. The encoder processes input sequences through a series of transformer blocks, each incorporating multi-head self-attention and position-wise feed-forward layers.",
    ],
    "Dataset Construction and Preprocessing": [
        "The primary dataset was compiled from three publicly available corpora spanning six typologically diverse languages: English, Mandarin Chinese, Arabic, Finnish, Swahili, and Quechua. Each corpus was selected to represent a different morphological typology.",
        "A team of twelve annotators, all holding graduate degrees in linguistics, performed the manual labeling. The annotation guidelines were developed iteratively over three pilot rounds, with inter-annotator agreement measured using both Cohen's kappa and Krippendorff's alpha.",
    ],
    "Experimental Setup": [
        "Five baseline models were selected for comparison: a standard BiLSTM with attention, a pre-trained multilingual BERT model, XLM-RoBERTa, a graph neural network variant, and a traditional feature-engineered SVM pipeline. Each baseline was trained with identical data splits and preprocessing.",
        "All experiments were conducted on a cluster of eight NVIDIA A100 GPUs with 80GB memory each. The software stack included PyTorch 2.1, HuggingFace Transformers 4.35, and custom training utilities developed for this project.",
    ],
    "Results and Analysis": [
        "The proposed approach achieved state-of-the-art results on four of the six evaluation benchmarks, with statistically significant improvements over the strongest baseline (p < 0.01, paired bootstrap test). The largest gains were observed for agglutinative languages, consistent with our hypothesis.",
        "The ablation study revealed that the hierarchical attention mechanism contributed the most to overall performance, accounting for approximately 3.2 percentage points of improvement. Removing the cross-lingual pre-training objective reduced accuracy by 2.8 points on average.",
    ],
    "Discussion": [
        "The findings presented in this dissertation carry several important implications for both the research community and practitioners working with low-resource languages. The demonstrated effectiveness of architecture search in this domain suggests that language-specific structural adaptations may be more important than previously recognized.",
        "From a theoretical perspective, the results challenge the prevailing assumption that a single universal architecture can serve all languages equally well. The observed correlation between typological features and optimal architecture choices opens new avenues for linguistically-informed model design.",
    ],
    "Conclusion and Future Work": [
        "This dissertation has made three primary contributions to the field of neural architecture search for natural language processing. First, we proposed a novel search strategy that incorporates typological priors into the architecture design space. Second, we demonstrated significant improvements for low-resource languages through targeted architectural adaptations.",
        "Future work will extend this research in several directions. The most immediate extension involves incorporating multi-modal inputs, combining text with speech and visual features. Additionally, we plan to explore the application of our search methodology to generative tasks such as machine translation and text summarization.",
    ],
}


def add_figure_caption(doc, fig_num, caption_text):
    """Add a caption paragraph for a figure using SEQ field codes."""
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_para.paragraph_format.space_before = Pt(6)
    caption_para.paragraph_format.space_after = Pt(12)

    # "Figure " label
    run_label = caption_para.add_run("Figure ")
    run_label.font.size = Pt(10)
    run_label.italic = True

    # SEQ field for figure number
    r_begin = caption_para.add_run()
    fld_begin = r_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r_begin._element.append(fld_begin)

    r_instr = caption_para.add_run()
    instr_text = r_instr._element.makeelement(qn('w:instrText'), {})
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' SEQ Figure \\* ARABIC '
    r_instr._element.append(instr_text)

    r_sep = caption_para.add_run()
    fld_sep = r_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r_sep._element.append(fld_sep)

    r_num = caption_para.add_run(str(fig_num))
    r_num.font.size = Pt(10)
    r_num.italic = True

    r_end = caption_para.add_run()
    fld_end = r_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r_end._element.append(fld_end)

    # ": Caption text"
    run_caption = caption_para.add_run(f": {caption_text}")
    run_caption.font.size = Pt(10)
    run_caption.italic = True

    # Mark this paragraph with the Caption style if available
    caption_para.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else doc.styles['Normal']


def add_table_caption(doc, table_num, caption_text):
    """Add a caption paragraph for a table using SEQ field codes."""
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_para.paragraph_format.space_before = Pt(12)
    caption_para.paragraph_format.space_after = Pt(6)

    run_label = caption_para.add_run("Table ")
    run_label.font.size = Pt(10)
    run_label.italic = True

    # SEQ field for table number
    r_begin = caption_para.add_run()
    fld_begin = r_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r_begin._element.append(fld_begin)

    r_instr = caption_para.add_run()
    instr_text = r_instr._element.makeelement(qn('w:instrText'), {})
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' SEQ Table \\* ARABIC '
    r_instr._element.append(instr_text)

    r_sep = caption_para.add_run()
    fld_sep = r_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r_sep._element.append(fld_sep)

    r_num = caption_para.add_run(str(table_num))
    r_num.font.size = Pt(10)
    r_num.italic = True

    r_end = caption_para.add_run()
    fld_end = r_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r_end._element.append(fld_end)

    run_caption = caption_para.add_run(f": {caption_text}")
    run_caption.font.size = Pt(10)
    run_caption.italic = True

    caption_para.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else doc.styles['Normal']


def add_abbreviation_entry(doc, abbr, full_form):
    """Add an abbreviation XE (index entry) field code for user-defined index."""
    # Create a hidden paragraph with XE field marking
    para = doc.add_paragraph()
    run_text = para.add_run(f"{abbr} ({full_form}) ")
    run_text.font.size = Pt(11)

    # Add XE field code: { XE "abbr:full_form" \f "Abbreviations" }
    r_begin = para.add_run()
    fld_begin = r_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r_begin._element.append(fld_begin)

    r_instr = para.add_run()
    instr_text = r_instr._element.makeelement(qn('w:instrText'), {})
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f' XE "{abbr}: {full_form}" \\f "Abbreviations" '
    r_instr._element.append(instr_text)

    r_sep = para.add_run()
    fld_sep = r_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r_sep._element.append(fld_sep)

    r_end = para.add_run()
    fld_end = r_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r_end._element.append(fld_end)


def create_sample_table(doc, num_rows, num_cols, headers, data_rows):
    """Create a formatted data table."""
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title Page ---
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading("Neural Architecture Search for Low-Resource Natural Language Processing", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("A Dissertation Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy")
    run.font.size = Pt(14)

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("\n\nElena Vasquez-Rodriguez")
    run.font.size = Pt(16)
    run.bold = True

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run("Department of Computer Science\nStanford University\nJune 2025")
    run.font.size = Pt(12)

    # Page break after title page
    doc.add_page_break()

    # Create placeholder image
    img_path = f'{WORKDIR}/placeholder_fig.png'
    create_placeholder_image(img_path)

    # Track figure and table counters
    fig_counter = 0
    table_counter = 0

    # --- Chapters ---
    for ch_idx, chapter in enumerate(CHAPTERS):
        # Chapter heading (Heading 1)
        ch_num = ch_idx + 1
        doc.add_heading(f"Chapter {ch_num}: {chapter['title']}", level=1)

        # Body text for chapter intro
        ch_title = chapter['title']
        if ch_title in BODY_TEXTS:
            for para_text in BODY_TEXTS[ch_title]:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)

        # Sub-sections
        for ss_idx, subsection in enumerate(chapter['subsections']):
            doc.add_heading(f"{ch_num}.{ss_idx + 1} {subsection}", level=2)

            # Add body text
            p = doc.add_paragraph(
                f"This section presents the {subsection.lower()} component of the research. "
                f"The analysis draws upon established methodologies in the field while introducing "
                f"novel adaptations suited to the specific challenges of low-resource language processing."
            )
            p.paragraph_format.space_after = Pt(6)

            # Add a second paragraph for substance
            p2 = doc.add_paragraph(
                f"The approach described here was developed through iterative refinement over "
                f"multiple experimental cycles, ensuring robustness across varying conditions "
                f"and dataset characteristics."
            )
            p2.paragraph_format.space_after = Pt(6)

        # Add figures for this chapter
        for fig_num in chapter.get('figures', []):
            fig_counter += 1
            # Add the image
            doc.add_picture(img_path, width=Inches(4.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add caption with SEQ field
            add_figure_caption(doc, fig_num, FIGURE_CAPTIONS[fig_num - 1])

        # Add tables for this chapter
        for tbl_num in chapter.get('tables', []):
            table_counter += 1
            # Table caption before table
            add_table_caption(doc, tbl_num, TABLE_CAPTIONS[tbl_num - 1])

            # Create sample data table
            if tbl_num == 1:
                create_sample_table(doc, 4, 3,
                    ["Research Question", "Method", "Chapter"],
                    [
                        ["RQ1: Architecture effectiveness", "Quantitative experiments", "5, 6"],
                        ["RQ2: Transfer learning impact", "Comparative analysis", "6"],
                        ["RQ3: Typological factors", "Correlation study", "6, 7"],
                    ])
            elif tbl_num == 2:
                create_sample_table(doc, 5, 4,
                    ["Approach", "Year", "Languages", "Performance"],
                    [
                        ["Zoph & Le", "2017", "1", "74.2%"],
                        ["Liu et al.", "2019", "3", "78.5%"],
                        ["Chen & Wu", "2021", "6", "81.3%"],
                        ["Park et al.", "2023", "12", "83.7%"],
                    ])
            elif tbl_num == 3:
                create_sample_table(doc, 5, 3,
                    ["Hyperparameter", "Search Range", "Selected Value"],
                    [
                        ["Learning rate", "1e-5 to 1e-3", "3e-4"],
                        ["Batch size", "16, 32, 64", "32"],
                        ["Hidden dimensions", "256, 512, 768", "512"],
                        ["Number of layers", "4, 6, 8, 12", "8"],
                    ])
            elif tbl_num == 4:
                create_sample_table(doc, 5, 4,
                    ["Language", "Train Samples", "Dev Samples", "Test Samples"],
                    [
                        ["English", "45,230", "5,640", "5,640"],
                        ["Mandarin", "38,710", "4,830", "4,830"],
                        ["Arabic", "22,450", "2,800", "2,800"],
                        ["Finnish", "15,680", "1,960", "1,960"],
                    ])
            elif tbl_num == 5:
                create_sample_table(doc, 4, 3,
                    ["Category", "Cohen's Kappa", "Krippendorff's Alpha"],
                    [
                        ["Entity Recognition", "0.87", "0.85"],
                        ["Relation Extraction", "0.79", "0.76"],
                        ["Sentiment Analysis", "0.91", "0.89"],
                    ])
            elif tbl_num == 6:
                create_sample_table(doc, 4, 3,
                    ["Component", "Specification", "Quantity"],
                    [
                        ["GPU", "NVIDIA A100 80GB", "8"],
                        ["CPU", "AMD EPYC 7763", "2"],
                        ["RAM", "512 GB DDR4", "1"],
                    ])
            elif tbl_num == 7:
                create_sample_table(doc, 6, 4,
                    ["Model", "Accuracy", "F1-Score", "BLEU"],
                    [
                        ["BiLSTM", "72.3%", "69.8", "28.4"],
                        ["mBERT", "78.1%", "75.6", "33.2"],
                        ["XLM-R", "80.5%", "78.2", "35.7"],
                        ["GNN Variant", "76.9%", "74.3", "31.5"],
                        ["Ours", "84.2%", "82.1", "39.8"],
                    ])
            elif tbl_num == 8:
                create_sample_table(doc, 5, 3,
                    ["Removed Component", "Accuracy Drop", "F1 Drop"],
                    [
                        ["Hierarchical attention", "-3.2%", "-3.5"],
                        ["Cross-lingual pretraining", "-2.8%", "-3.0"],
                        ["Typological features", "-1.5%", "-1.8"],
                        ["Data augmentation", "-1.1%", "-1.3"],
                    ])
            elif tbl_num == 9:
                create_sample_table(doc, 4, 3,
                    ["Requirement", "Specification", "Status"],
                    [
                        ["Inference time", "< 100ms per sample", "Met"],
                        ["Memory footprint", "< 2GB GPU", "Met"],
                        ["API throughput", "> 500 req/sec", "Pending"],
                    ])
            elif tbl_num == 10:
                create_sample_table(doc, 4, 3,
                    ["Contribution", "Description", "Chapter"],
                    [
                        ["NAS with typological priors", "Novel search strategy", "3, 5, 6"],
                        ["Low-resource improvements", "Targeted adaptations", "4, 6"],
                        ["Multi-lingual benchmark", "New evaluation suite", "4, 6, 7"],
                    ])

            doc.add_paragraph()  # spacing after table

    # --- Abbreviation entries scattered throughout (marked with XE fields) ---
    # Add a section at the end with abbreviation entries marked for indexing
    doc.add_heading("Glossary of Terms", level=1)
    p = doc.add_paragraph(
        "The following abbreviations are used throughout this dissertation. "
        "Each term is defined at its first occurrence in the text."
    )
    p.paragraph_format.space_after = Pt(6)

    for abbr, full_form in ABBREVIATIONS:
        add_abbreviation_entry(doc, abbr, full_form)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up placeholder image
    try:
        os.remove(img_path)
    except OSError:
        pass

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
