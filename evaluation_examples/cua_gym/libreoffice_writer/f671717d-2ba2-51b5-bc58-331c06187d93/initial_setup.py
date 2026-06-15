"""
Initial Setup: Court filing document with single line spacing
Task ID: writer_legal_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_paragraph(doc, text, level=1):
    """Add a heading with explicit single spacing kept at default."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_body_paragraph(doc, text, bold=False, alignment=None):
    """Add a body paragraph with explicit single (1.0) line spacing."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    if alignment:
        para.paragraph_format.alignment = alignment
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    return para


def create_initial():
    doc = Document()

    # Page setup - standard legal formatting
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Court Caption ---
    add_body_paragraph(
        doc,
        "UNITED STATES DISTRICT COURT",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    )
    add_body_paragraph(
        doc,
        "NORTHERN DISTRICT OF CALIFORNIA",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    )
    add_body_paragraph(
        doc,
        "SAN FRANCISCO DIVISION",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    # Blank line
    add_body_paragraph(doc, "")

    # Case parties
    add_body_paragraph(doc, "GREENLEAF TECHNOLOGIES, INC.,")
    add_body_paragraph(doc, "    Plaintiff,")
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "v.", alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "BRIGHTON ANALYTICS GROUP, LLC,")
    add_body_paragraph(doc, "    Defendant.")

    # Case number
    add_body_paragraph(doc, "")
    add_body_paragraph(
        doc,
        "Case No. 3:2025-cv-04187-WHJ",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
    )
    add_body_paragraph(doc, "")

    # --- Title ---
    add_heading_paragraph(doc, "DEFENDANT'S MOTION TO DISMISS PURSUANT TO RULE 12(b)(6)", level=1)

    # --- Introduction ---
    add_heading_paragraph(doc, "I. INTRODUCTION", level=2)

    add_body_paragraph(
        doc,
        "Defendant Brighton Analytics Group, LLC ('Brighton' or 'Defendant'), by and through "
        "its undersigned counsel, hereby moves this Court to dismiss the Complaint filed by "
        "Plaintiff Greenleaf Technologies, Inc. ('Greenleaf' or 'Plaintiff') pursuant to "
        "Federal Rule of Civil Procedure 12(b)(6) for failure to state a claim upon which "
        "relief can be granted."
    )

    add_body_paragraph(
        doc,
        "As set forth more fully below, Plaintiff's Complaint is fatally deficient in several "
        "respects. The allegations contained therein are conclusory, fail to establish the "
        "essential elements of the claims asserted, and do not meet the plausibility standard "
        "established by the Supreme Court in Bell Atlantic Corp. v. Twombly, 550 U.S. 544 "
        "(2007) and Ashcroft v. Iqbal, 556 U.S. 662 (2009)."
    )

    # --- Factual Background ---
    add_heading_paragraph(doc, "II. FACTUAL BACKGROUND", level=2)

    add_body_paragraph(
        doc,
        "On or about January 15, 2025, Plaintiff and Defendant entered into a Software "
        "Licensing Agreement ('the Agreement') whereby Defendant licensed certain proprietary "
        "data analytics software to Plaintiff for use in Plaintiff's internal business "
        "operations. The Agreement was executed at Defendant's principal office in Palo Alto, "
        "California."
    )

    add_body_paragraph(
        doc,
        "Pursuant to Section 4.2 of the Agreement, Plaintiff was granted a non-exclusive, "
        "non-transferable license to use the BrightView Analytics Platform ('the Software') "
        "for a term of twenty-four (24) months commencing on February 1, 2025. The total "
        "licensing fee was $487,500, payable in quarterly installments of $60,937.50."
    )

    add_body_paragraph(
        doc,
        "On March 28, 2025, Plaintiff transmitted a letter to Defendant alleging that the "
        "Software contained material defects that rendered it unsuitable for its intended "
        "purpose. Specifically, Plaintiff claimed that the predictive modeling module produced "
        "results with an error rate exceeding 12%, whereas the Agreement warranted accuracy "
        "within a 3% margin of error."
    )

    add_body_paragraph(
        doc,
        "Defendant responded on April 10, 2025, noting that its engineering team had conducted "
        "an independent audit of the Software and found no defects. Defendant further noted "
        "that Plaintiff's reported errors were attributable to improper data formatting on "
        "Plaintiff's end, which violated Section 7.1 of the Agreement's technical "
        "specifications."
    )

    # --- Legal Standard ---
    add_heading_paragraph(doc, "III. LEGAL STANDARD", level=2)

    add_body_paragraph(
        doc,
        "Under Federal Rule of Civil Procedure 12(b)(6), a court must dismiss a complaint "
        "if it fails to state a claim upon which relief can be granted. To survive a motion "
        "to dismiss, a complaint must contain 'enough facts to state a claim to relief that "
        "is plausible on its face.' Bell Atlantic Corp. v. Twombly, 550 U.S. 544, 570 (2007)."
    )

    add_body_paragraph(
        doc,
        "A claim is facially plausible 'when the plaintiff pleads factual content that allows "
        "the court to draw the reasonable inference that the defendant is liable for the "
        "misconduct alleged.' Ashcroft v. Iqbal, 556 U.S. 662, 678 (2009). While the court "
        "must accept all well-pleaded factual allegations as true, it need not accept legal "
        "conclusions cast as factual allegations. Id. at 679."
    )

    # --- Argument ---
    add_heading_paragraph(doc, "IV. ARGUMENT", level=2)

    add_heading_paragraph(
        doc,
        "A. Plaintiff's Breach of Contract Claim Fails Because the Complaint Does Not "
        "Adequately Allege the Existence of a Material Breach",
        level=3
    )

    add_body_paragraph(
        doc,
        "To state a claim for breach of contract under California law, a plaintiff must allege: "
        "(1) the existence of a valid contract; (2) the plaintiff's performance or excuse for "
        "nonperformance; (3) the defendant's material breach; and (4) resulting damages. "
        "Oasis West Realty, LLC v. Goldman, 51 Cal. 4th 811, 821 (2011)."
    )

    add_body_paragraph(
        doc,
        "Here, Plaintiff's Complaint fails to adequately allege the third element. The "
        "Complaint merely states in conclusory fashion that 'the Software did not perform as "
        "warranted.' This bare assertion, devoid of supporting factual detail, is precisely "
        "the type of formulaic recitation of elements that Twombly and Iqbal prohibit."
    )

    add_body_paragraph(
        doc,
        "Moreover, Section 9.3 of the Agreement contains a limitation of liability clause "
        "that expressly excludes consequential damages arising from software performance "
        "issues. Plaintiff's Complaint seeks $2.3 million in lost revenue projections, which "
        "constitute precisely the type of consequential damages barred by this provision."
    )

    add_heading_paragraph(
        doc,
        "B. Plaintiff's Negligent Misrepresentation Claim Is Barred by the Economic Loss Rule",
        level=3
    )

    add_body_paragraph(
        doc,
        "California courts have consistently held that the economic loss rule bars tort claims "
        "that arise from the same operative facts as a breach of contract claim. Robinson "
        "Helicopter Co., Inc. v. Dana Corp., 34 Cal. 4th 979, 988 (2004). Plaintiff's "
        "negligent misrepresentation claim is predicated entirely on Defendant's alleged "
        "representations regarding software performance, which are the same representations "
        "embodied in the Agreement."
    )

    add_body_paragraph(
        doc,
        "Because Plaintiff has not alleged any duty independent of the contractual obligations, "
        "or any injury beyond the economic losses flowing from the alleged breach, the "
        "negligent misrepresentation claim must be dismissed as a matter of law."
    )

    # --- Conclusion ---
    add_heading_paragraph(doc, "V. CONCLUSION", level=2)

    add_body_paragraph(
        doc,
        "For the foregoing reasons, Defendant Brighton Analytics Group, LLC respectfully "
        "requests that this Court grant its Motion to Dismiss and dismiss Plaintiff's "
        "Complaint in its entirety with prejudice. In the alternative, Defendant requests "
        "that the Court dismiss the negligent misrepresentation claim with prejudice and the "
        "breach of contract claim without prejudice, with leave to amend."
    )

    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "Respectfully submitted,")
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "HARRISON & CROSS LLP", bold=True)
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "By: _________________________________")
    add_body_paragraph(doc, "    Rebecca A. Harrison, Esq.")
    add_body_paragraph(doc, "    California Bar No. 287451")
    add_body_paragraph(doc, "    525 Market Street, Suite 3200")
    add_body_paragraph(doc, "    San Francisco, CA 94105")
    add_body_paragraph(doc, "    Tel: (415) 555-0142")
    add_body_paragraph(doc, "    Email: rharrison@harrisoncross.com")
    add_body_paragraph(doc, "")
    add_body_paragraph(doc, "    Attorneys for Defendant")
    add_body_paragraph(doc, "    Brighton Analytics Group, LLC")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
