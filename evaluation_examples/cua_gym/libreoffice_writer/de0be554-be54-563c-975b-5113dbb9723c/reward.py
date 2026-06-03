"""
Reward Script: Insert a 3-column, 4-row table below the second paragraph
Task ID: osworld_writer_table_creation_001
Domain: libreoffice_writer
Scoring:
  Component 1: Table exists in document (0.4 pts)
  Component 2: Table has correct dimensions — 4 rows x 3 columns (0.4 pts)
  Component 3: Table is positioned after the 2nd paragraph (0.2 pts)
"""

import os

from docx import Document

WORKDIR = '/home/user'  # VM path — all reward scripts run on the VM
TASK_ID = 'osworld_writer_table_creation_001'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.

    Task: Insert a 3-column, 4-row table below the second paragraph of a
    3-paragraph business letter that initially has no tables.

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document — precondition gate
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: A table has been inserted into the document (0.4 points)
    # Initial state has 0 tables; golden state should have 1 table.
    try:
        num_tables = len(doc.tables)
        if num_tables >= 1:
            print(f"PASS: Component 1 — At least one table found (count={num_tables}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Expected at least 1 table, found {num_tables}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table has correct dimensions — 4 rows and 3 columns (0.4 points)
    # Task specifies a 3-column, 4-row table explicitly.
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 4 and num_cols == 3:
                print(f"PASS: Component 2 — Table dimensions correct: {num_rows} rows x {num_cols} cols (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — Expected 4 rows x 3 cols, found {num_rows} rows x {num_cols} cols")
        else:
            print("FAIL: Component 2 — No table present to check dimensions")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table is positioned after the 2nd paragraph (0.2 points)
    # The document body should have: paragraph[0], paragraph[1], table, paragraph[2], ...
    # We verify the table appears as the 3rd child element in the document body
    # (index 2), immediately following the 2nd paragraph (index 1).
    try:
        body = doc.element.body
        children = list(body)
        # Strip the sectPr element (last element is section properties, not content)
        content_children = [c for c in children if c.tag.split('}')[-1] != 'sectPr']

        # Find the position of the first table element
        table_index = None
        for idx, child in enumerate(content_children):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'tbl':
                table_index = idx
                break

        if table_index is None:
            print("FAIL: Component 3 — No table found in body to check position")
        elif table_index == 2:
            # Index 2 means: [0]=para1, [1]=para2, [2]=table — correct position
            print(f"PASS: Component 3 — Table is at body index {table_index}, correctly after 2nd paragraph (0.2 pts)")
            total_score += 0.2
        else:
            # Describe actual position
            before_table = table_index  # number of paragraphs before the table
            print(f"FAIL: Component 3 — Table is at body index {table_index}, "
                  f"expected index 2 (after 2nd paragraph). "
                  f"Table appears after {before_table} element(s) instead of 2.")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
