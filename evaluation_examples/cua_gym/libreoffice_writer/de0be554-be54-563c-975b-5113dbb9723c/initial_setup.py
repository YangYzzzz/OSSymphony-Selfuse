"""
Initial Setup: Insert a 3-column, 4-row table below the second paragraph of this document.
Task ID: osworld_writer_table_creation_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_table_creation_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Business Letter: 3 paragraphs, no tables ---

    # Paragraph 1: Date and salutation
    p1 = doc.add_paragraph(
        "March 5, 2025\n\nDear Mr. Thompson,"
    )
    p1.paragraph_format.space_after = Pt(12)

    # Paragraph 2: Body paragraph 1 (opening)
    p2 = doc.add_paragraph(
        "I am writing to follow up on our recent conversation regarding the Q1 budget "
        "allocation for the Northern Region. As discussed during our meeting on February 28, "
        "we have identified several areas where operational efficiency can be improved, "
        "particularly in the logistics and procurement departments. Our team has conducted "
        "a thorough analysis of expenditure patterns over the past fiscal year and is ready "
        "to present our findings."
    )
    p2.paragraph_format.space_after = Pt(12)

    # Paragraph 3: Body paragraph 2 (closing)
    p3 = doc.add_paragraph(
        "We would appreciate the opportunity to schedule a follow-up meeting at your earliest "
        "convenience to review the detailed report. Please feel free to contact our project "
        "coordinator, Ms. Diana Reyes, at extension 4821, should you require any preliminary "
        "information before the meeting. We look forward to your response and to continuing "
        "our productive collaboration.\n\n"
        "Sincerely,\n\nJonathan Hargrove\nRegional Manager, Northern Division"
    )
    p3.paragraph_format.space_after = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
