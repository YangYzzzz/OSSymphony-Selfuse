"""
Initial Setup: Emma Birthday Party Invitation
Task ID: writer_creative_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_003'
OUTPUT = f'{WORKDIR}/emma_birthday_invite.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default empty paragraph if present
    # Add Line 1: Title - 12pt, left-aligned, regular (NOT bold), black, Times New Roman
    para1 = doc.add_paragraph()
    para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run1 = para1.add_run("You\u2019re Invited to Emma\u2019s Birthday Party!")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(12)
    run1.bold = False
    run1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add Line 2: Date detail - 12pt, left-aligned, NOT bold
    para2 = doc.add_paragraph()
    para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run2 = para2.add_run("Date: March 14, 2026")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(12)
    run2.bold = False
    run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add Line 3: Time detail - 12pt, left-aligned, NOT bold
    para3 = doc.add_paragraph()
    para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run3 = para3.add_run("Time: 2:00 PM - 5:00 PM")
    run3.font.name = "Times New Roman"
    run3.font.size = Pt(12)
    run3.bold = False
    run3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add Line 4: Place detail - 12pt, left-aligned, NOT bold
    para4 = doc.add_paragraph()
    para4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run4 = para4.add_run("Place: Sunshine Park Pavilion")
    run4.font.name = "Times New Roman"
    run4.font.size = Pt(12)
    run4.bold = False
    run4.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add Line 5: RSVP detail - 12pt, left-aligned, NOT bold
    para5 = doc.add_paragraph()
    para5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run5 = para5.add_run("RSVP: (503) 555-0234")
    run5.font.name = "Times New Roman"
    run5.font.size = Pt(12)
    run5.bold = False
    run5.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add Line 6: Closing line - 12pt, left-aligned, NOT bold
    para6 = doc.add_paragraph()
    para6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run6 = para6.add_run("Come join us for cake, games, and fun!")
    run6.font.name = "Times New Roman"
    run6.font.size = Pt(12)
    run6.bold = False
    run6.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
