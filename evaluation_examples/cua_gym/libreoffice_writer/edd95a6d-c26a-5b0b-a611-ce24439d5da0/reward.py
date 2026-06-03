"""
Reward Script: Birthday Party Invitation Formatting
Task ID: writer_creative_003
Domain: libreoffice_writer
Scoring:
  Component 1: Title paragraph — 28pt, bold, centered (0.4 pts)
  Component 2: Title color is hot pink #FF69B4 (0.2 pts)
  Component 3: Detail lines (Date/Time/Place/RSVP) — 14pt, centered, label bold (0.3 pts)
  Component 4: Footer line "Come join us..." — 14pt, centered (0.1 pts)
  Total: 1.0
"""

import os
from math import sqrt

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_003'
FILE_PATH = f'{WORKDIR}/emma_birthday_invite.docx'

# Color comparison helper — Euclidean RGB distance
def color_distance(c1, c2):
    """Compute Euclidean RGB distance between two (R, G, B) tuples."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))

HOT_PINK = (0xFF, 0x69, 0xB4)
COLOR_TOLERANCE = 30  # Euclidean distance threshold for color match


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Sanity check: document should have 6 paragraphs
    if len(doc.paragraphs) < 6:
        print(f"CRITICAL: Expected at least 6 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    para_title = doc.paragraphs[0]
    para_date = doc.paragraphs[1]
    para_time = doc.paragraphs[2]
    para_place = doc.paragraphs[3]
    para_rsvp = doc.paragraphs[4]
    para_footer = doc.paragraphs[5]

    # -----------------------------------------------------------------------
    # Component 1: Title paragraph — 28pt, bold, centered alignment (0.4 pts)
    # -----------------------------------------------------------------------
    try:
        title_text = para_title.text.strip()

        # Check alignment is CENTER
        title_alignment = para_title.paragraph_format.alignment
        is_centered = (title_alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

        # Check all runs in title for size=28pt and bold=True
        title_runs = [r for r in para_title.runs if r.text.strip()]
        all_28pt = all(
            r.font.size is not None and abs(r.font.size.pt - 28.0) < 0.5
            for r in title_runs
        )
        all_bold = all(r.font.bold is True for r in title_runs)

        if is_centered and all_28pt and all_bold and title_runs:
            print(f"PASS: Component 1 — Title is 28pt, bold, centered (alignment={title_alignment}, runs={len(title_runs)}) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Title formatting: centered={is_centered}, all_28pt={all_28pt}, all_bold={all_bold}, title_text={repr(title_text[:40])}")
            if title_runs:
                sample = title_runs[0]
                print(f"  Sample run: size={sample.font.size.pt if sample.font.size else None}, bold={sample.font.bold}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Title color is hot pink (#FF69B4) (0.2 pts)
    # -----------------------------------------------------------------------
    try:
        title_runs = [r for r in para_title.runs if r.text.strip()]
        colored_runs = []
        for r in title_runs:
            try:
                if r.font.color and r.font.color.rgb is not None:
                    rgb = r.font.color.rgb
                    colored_runs.append((r.text, (rgb[0], rgb[1], rgb[2])))
            except Exception:
                pass

        if not colored_runs:
            print(f"FAIL: Component 2 — No title runs have an explicit color set")
        else:
            all_pink = all(
                color_distance(c, HOT_PINK) < COLOR_TOLERANCE
                for _, c in colored_runs
            )
            if all_pink:
                print(f"PASS: Component 2 — Title is hot pink (#FF69B4), runs checked: {len(colored_runs)} (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 2 — Title color mismatch. Expected ~#FF69B4, found: {colored_runs}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Detail lines (Date/Time/Place/RSVP) — 14pt, centered,
    #              label portion bold (0.3 pts)
    # -----------------------------------------------------------------------
    try:
        detail_paras = [para_date, para_time, para_place, para_rsvp]
        detail_labels = ['Date:', 'Time:', 'Place:', 'RSVP:']

        all_detail_ok = True
        detail_issues = []

        for para, label in zip(detail_paras, detail_labels):
            # Check alignment is CENTER
            align = para.paragraph_format.alignment
            if align != WD_PARAGRAPH_ALIGNMENT.CENTER:
                all_detail_ok = False
                detail_issues.append(f"{label} alignment={align} (expected CENTER)")

            # Check all runs are 14pt
            runs = [r for r in para.runs if r.text.strip()]
            for r in runs:
                size_pt = r.font.size.pt if r.font.size else None
                if size_pt is None or abs(size_pt - 14.0) > 0.5:
                    all_detail_ok = False
                    detail_issues.append(f"{label} run '{r.text[:20]}' size={size_pt} (expected 14pt)")

            # Check that the label run is bold
            label_run_found = False
            for r in para.runs:
                if r.text.strip().startswith(label.rstrip(':')) or r.text.strip() == label:
                    label_run_found = True
                    if not r.font.bold:
                        all_detail_ok = False
                        detail_issues.append(f"{label} label run bold={r.font.bold} (expected True)")
                    break
            if not label_run_found and runs:
                # Also check if the first run contains the label and is bold
                first_run = runs[0]
                if label.split(':')[0] in first_run.text:
                    label_run_found = True
                    if not first_run.font.bold:
                        all_detail_ok = False
                        detail_issues.append(f"{label} first run bold={first_run.font.bold} (expected True)")
            if not label_run_found:
                all_detail_ok = False
                detail_issues.append(f"{label} label run not found")

        if all_detail_ok:
            print(f"PASS: Component 3 — All 4 detail lines are 14pt, centered, labels bold (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Detail lines issues: {detail_issues}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Footer line "Come join us..." — 14pt, centered (0.1 pts)
    # -----------------------------------------------------------------------
    try:
        footer_align = para_footer.paragraph_format.alignment
        footer_runs = [r for r in para_footer.runs if r.text.strip()]

        footer_centered = (footer_align == WD_PARAGRAPH_ALIGNMENT.CENTER)
        footer_14pt = all(
            r.font.size is not None and abs(r.font.size.pt - 14.0) < 0.5
            for r in footer_runs
        )

        if footer_centered and footer_14pt and footer_runs:
            print(f"PASS: Component 4 — Footer line is 14pt, centered (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 4 — Footer line: centered={footer_centered}, all_14pt={footer_14pt}")
            if footer_runs:
                sample = footer_runs[0]
                print(f"  Sample run: size={sample.font.size.pt if sample.font.size else None}, alignment={footer_align}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    # Final score
    # -----------------------------------------------------------------------
    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
