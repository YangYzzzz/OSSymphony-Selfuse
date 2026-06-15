"""
Initial Setup: API documentation document with sample content using default styles.
Task ID: writer_tech_055
Domain: libreoffice_writer

Creates a Writer document with realistic API documentation content but using only
default/built-in styles. No custom paragraph or character styles are defined.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # ---- Page Setup ----
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ============================================================
    # COVER PAGE (simple, using default styles)
    # ============================================================
    # Title area
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    title_run = title_para.add_run("CloudSync Platform")
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle_para.add_run("REST API Reference Guide")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # Version info
    ver_para = doc.add_paragraph()
    ver_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ver_para.paragraph_format.space_before = Pt(48)
    ver_run = ver_para.add_run("Version 3.2.1")
    ver_run.font.size = Pt(12)
    ver_run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run("Last Updated: March 2026")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    # Page break after cover
    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS (placeholder)
    # ============================================================
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Authentication",
        "3. Users API",
        "4. Projects API",
        "5. Files API",
        "6. Error Handling",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ============================================================
    # SECTION 1: Introduction
    # ============================================================
    doc.add_heading("1. Introduction", level=1)

    doc.add_paragraph(
        "The CloudSync Platform API provides programmatic access to cloud storage, "
        "file synchronization, and team collaboration features. This document describes "
        "the available REST endpoints, request/response formats, authentication mechanisms, "
        "and error handling procedures."
    )

    doc.add_paragraph(
        "All API requests must be made over HTTPS. Calls made over plain HTTP will be "
        "rejected. The base URL for all API endpoints is:"
    )

    # Code example using plain paragraph (no custom style)
    code_para = doc.add_paragraph("https://api.cloudsync.io/v3")
    code_run = code_para.runs[0]
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(10)

    # ============================================================
    # SECTION 2: Authentication
    # ============================================================
    doc.add_heading("2. Authentication", level=1)

    doc.add_paragraph(
        "The CloudSync API uses OAuth 2.0 Bearer tokens for authentication. "
        "Include your access token in the Authorization header of each request."
    )

    doc.add_heading("Obtaining an Access Token", level=2)

    doc.add_paragraph(
        "To obtain an access token, make a POST request to the token endpoint "
        "with your client credentials:"
    )

    # Code block (plain formatting, no custom style)
    code_lines = [
        'POST /oauth/token HTTP/1.1',
        'Host: auth.cloudsync.io',
        'Content-Type: application/x-www-form-urlencoded',
        '',
        'grant_type=client_credentials',
        '&client_id=cs_live_a8f3k2m9x1',
        '&client_secret=sk_prod_7hG3nR9pL2mK',
    ]
    for line in code_lines:
        p = doc.add_paragraph(line if line else " ")
        for r in p.runs:
            r.font.name = "Courier New"
            r.font.size = Pt(10)

    # Note (plain paragraph)
    note_para = doc.add_paragraph()
    note_run = note_para.add_run("Note: ")
    note_run.bold = True
    note_para.add_run(
        "Access tokens expire after 3600 seconds. Use the refresh_token grant type "
        "to obtain a new token without requiring user re-authentication."
    )

    # ============================================================
    # SECTION 3: Users API
    # ============================================================
    doc.add_heading("3. Users API", level=1)

    doc.add_heading("GET /users/{user_id}", level=2)

    doc.add_paragraph(
        "Retrieves detailed information about a specific user account, including "
        "profile data, subscription tier, and storage usage statistics."
    )

    # Parameter table (using default Table Grid style)
    doc.add_heading("Parameters", level=3)
    param_table = doc.add_table(rows=4, cols=4)
    param_table.style = "Table Grid"

    headers = ["Parameter", "Type", "Required", "Description"]
    for i, h in enumerate(headers):
        cell = param_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    params = [
        ["user_id", "string", "Yes", "Unique identifier for the user (UUID format)"],
        ["include_stats", "boolean", "No", "Include storage usage statistics in response"],
        ["fields", "string", "No", "Comma-separated list of fields to return"],
    ]
    for r, row_data in enumerate(params, 1):
        for c, val in enumerate(row_data):
            param_table.cell(r, c).text = val

    doc.add_heading("Example Response", level=3)

    response_lines = [
        '{',
        '  "user_id": "usr_8f3a2k9m1x7p",',
        '  "email": "sarah.chen@techcorp.io",',
        '  "display_name": "Sarah Chen",',
        '  "subscription": "enterprise",',
        '  "storage_used_gb": 142.7,',
        '  "storage_limit_gb": 500,',
        '  "created_at": "2024-08-15T09:23:41Z"',
        '}',
    ]
    for line in response_lines:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = "Courier New"
            r.font.size = Pt(10)

    # ============================================================
    # SECTION 4: Projects API
    # ============================================================
    doc.add_heading("4. Projects API", level=1)

    doc.add_heading("POST /projects", level=2)

    doc.add_paragraph(
        "Creates a new project workspace with the specified configuration. "
        "Projects serve as top-level containers for organizing files and "
        "managing team access permissions."
    )

    doc.add_heading("Request Body", level=3)
    body_table = doc.add_table(rows=5, cols=4)
    body_table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = body_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    body_params = [
        ["name", "string", "Yes", "Project display name (max 128 characters)"],
        ["description", "string", "No", "Brief project description (max 500 characters)"],
        ["visibility", "string", "No", "Access level: private, team, or public (default: private)"],
        ["team_id", "string", "No", "Team identifier for shared project access"],
    ]
    for r, row_data in enumerate(body_params, 1):
        for c, val in enumerate(row_data):
            body_table.cell(r, c).text = val

    # Note
    note2 = doc.add_paragraph()
    note2_run = note2.add_run("Important: ")
    note2_run.bold = True
    note2.add_run(
        "Free-tier accounts are limited to 5 active projects. Enterprise accounts "
        "have no project limit. Attempting to exceed the limit returns HTTP 429."
    )

    # ============================================================
    # SECTION 5: Files API
    # ============================================================
    doc.add_heading("5. Files API", level=1)

    doc.add_heading("PUT /projects/{project_id}/files/{file_path}", level=2)

    doc.add_paragraph(
        "Uploads or updates a file within a project. Supports multipart uploads "
        "for files exceeding 100 MB. The API automatically handles deduplication "
        "and versioning for all file operations."
    )

    doc.add_heading("Headers", level=3)
    hdr_table = doc.add_table(rows=4, cols=3)
    hdr_table.style = "Table Grid"

    hdr_headers = ["Header", "Value", "Description"]
    for i, h in enumerate(hdr_headers):
        cell = hdr_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    hdr_data = [
        ["Content-Type", "multipart/form-data", "Required for file uploads"],
        ["X-CS-Checksum", "sha256:<hash>", "Optional integrity verification"],
        ["X-CS-Conflict", "overwrite | rename | fail", "Conflict resolution strategy"],
    ]
    for r, row_data in enumerate(hdr_data, 1):
        for c, val in enumerate(row_data):
            hdr_table.cell(r, c).text = val

    # ============================================================
    # SECTION 6: Error Handling
    # ============================================================
    doc.add_heading("6. Error Handling", level=1)

    doc.add_paragraph(
        "The CloudSync API uses standard HTTP status codes to indicate the success "
        "or failure of an API request. Error responses include a JSON body with "
        "details about the specific error condition."
    )

    err_table = doc.add_table(rows=7, cols=3)
    err_table.style = "Table Grid"

    err_headers = ["Status Code", "Error Type", "Description"]
    for i, h in enumerate(err_headers):
        cell = err_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    errors = [
        ["400", "bad_request", "The request body contains invalid parameters"],
        ["401", "unauthorized", "Missing or invalid authentication token"],
        ["403", "forbidden", "Insufficient permissions for the requested resource"],
        ["404", "not_found", "The specified resource does not exist"],
        ["429", "rate_limited", "Too many requests; retry after Retry-After seconds"],
        ["500", "internal_error", "An unexpected server error occurred"],
    ]
    for r, row_data in enumerate(errors, 1):
        for c, val in enumerate(row_data):
            err_table.cell(r, c).text = val

    # Error response example
    doc.add_heading("Error Response Format", level=2)
    err_response = [
        '{',
        '  "error": {',
        '    "type": "unauthorized",',
        '    "message": "The provided access token has expired.",',
        '    "request_id": "req_3k8m2f9a1x7p",',
        '    "documentation_url": "https://docs.cloudsync.io/errors/unauthorized"',
        '  }',
        '}',
    ]
    for line in err_response:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = "Courier New"
            r.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
