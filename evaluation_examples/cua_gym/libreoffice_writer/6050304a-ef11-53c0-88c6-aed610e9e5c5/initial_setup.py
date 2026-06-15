"""
Initial Setup: complaint_letter.docx with default paragraph spacing
Task ID: writer_para_074
Domain: libreoffice_writer

Creates a complaint letter with 11 paragraphs:
- Paragraphs 1-5: address block (no explicit spacing set - default)
- Paragraph 6: date
- Paragraphs 7-9: body paragraphs (no explicit spacing set - default)
- Paragraphs 10-11: closing
The task is to set spacing for address block and body paragraphs.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_para_074'
OUTPUT = f'{WORKDIR}/Desktop/complaint_letter.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document()

    # Remove default paragraph spacing from Normal style
    # (leave as-is to simulate a typical document without forced spacing)

    # Paragraph 1: Robert Chen
    p1 = doc.add_paragraph('Robert Chen')

    # Paragraph 2: address line 1
    p2 = doc.add_paragraph('789 Elm Street')

    # Paragraph 3: address line 2
    p3 = doc.add_paragraph('Chicago, IL 60601')

    # Paragraph 4: phone
    p4 = doc.add_paragraph('Phone: (312) 555-0987')

    # Paragraph 5: email
    p5 = doc.add_paragraph('Email: r.chen@email.com')

    # Paragraph 6: date
    p6 = doc.add_paragraph('March 3, 2025')

    # Paragraph 7: first body paragraph
    p7 = doc.add_paragraph(
        'I am writing to formally complain about the defective refrigerator model FR-2000 '
        'that I purchased from your store on February 1, 2025. Despite following all '
        'installation instructions, the unit has failed to maintain proper temperature since delivery.'
    )

    # Paragraph 8: second body paragraph
    p8 = doc.add_paragraph(
        'I have contacted your customer service department three times and was promised a '
        'technician visit that has not materialized. This situation is unacceptable given that '
        'I paid $2,400 for what was advertised as a premium appliance.'
    )

    # Paragraph 9: third body paragraph
    p9 = doc.add_paragraph(
        'I request either a full replacement unit or a complete refund within 14 business days. '
        'Failure to resolve this matter will compel me to pursue remedies through consumer protection agencies.'
    )

    # Paragraph 10: closing
    p10 = doc.add_paragraph('Regards,')

    # Paragraph 11: signature
    p11 = doc.add_paragraph('Robert Chen')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
