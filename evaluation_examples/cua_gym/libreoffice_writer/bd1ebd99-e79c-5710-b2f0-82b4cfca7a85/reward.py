"""
Reward Script: Create an exam paper in LibreOffice Writer
Task ID: writer_wf_011
Domain: libreoffice_writer
Scoring:
  Component 1: Header with course name and exam date (0.15)
  Component 2: Instructions section with time limit, materials, grading (0.15)
  Component 3: Part A - Multiple Choice with 5 questions and a-d options (0.20)
  Component 4: Part B - Short Answer with 3 questions (0.15)
  Component 5: Part C - Essay with 1 question (0.10)
  Component 6: Page breaks before each part (0.15)
  Component 7: Answer key for Part A on separate page (0.10)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_011'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def find_heading_para_index(doc, pattern):
    """Find the index of a heading paragraph matching a pattern (case-insensitive).
    Prefers Heading-styled paragraphs, falls back to any matching paragraph."""
    # First try: heading-styled paragraphs
    for i, p in enumerate(doc.paragraphs):
        if p.style and 'heading' in p.style.name.lower():
            if re.search(pattern, p.text.lower()):
                return i
    # Fallback: any paragraph starting with the pattern
    for i, p in enumerate(doc.paragraphs):
        if re.match(r'\s*' + pattern, p.text.strip().lower()):
            return i
    return -1


def get_section_text(doc, start_idx, end_idx):
    """Get concatenated text of paragraphs between start_idx and end_idx (exclusive)."""
    texts = []
    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        texts.append(doc.paragraphs[i].text)
    return '\n'.join(texts)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_paras = len(doc.paragraphs)
    if num_paras == 0:
        print("FAIL: Document is empty (0 paragraphs)")
        print("REWARD: 0.0")
        return 0.0

    all_text = '\n'.join(p.text for p in doc.paragraphs)
    all_text_lower = all_text.lower()

    # Locate key section heading indices
    idx_part_a = find_heading_para_index(doc, r'part\s*a\s*[:\-–—]?\s*multiple\s*choice')
    idx_part_b = find_heading_para_index(doc, r'part\s*b\s*[:\-–—]?\s*short\s*answer')
    idx_part_c = find_heading_para_index(doc, r'part\s*c\s*[:\-–—]?\s*essay')
    idx_answer_key = find_heading_para_index(doc, r'answer\s*key')

    # Component 1: Header with course name and exam date (0.15 points)
    try:
        has_course_name = 'introduction to economics' in all_text_lower
        has_final_exam = 'final exam' in all_text_lower
        has_exam_date = bool(re.search(r'(exam\s*date|date\s*:?\s*\w+\s+\d)', all_text_lower))

        if has_course_name and has_final_exam and has_exam_date:
            print(f"PASS: Component 1 — Header has course name, 'final exam', and exam date (0.15 pts)")
            total_score += 0.15
        else:
            missing = []
            if not has_course_name:
                missing.append("course name 'Introduction to Economics'")
            if not has_final_exam:
                missing.append("'Final Exam'")
            if not has_exam_date:
                missing.append("exam date")
            print(f"FAIL: Component 1 — Missing: {', '.join(missing)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Instructions section with time limit, allowed materials, grading policy (0.15 points)
    try:
        has_time_limit = bool(re.search(r'time\s*limit', all_text_lower))
        has_materials = bool(re.search(r'allowed\s*materials?', all_text_lower))
        has_grading = bool(re.search(r'grading\s*(policy|:)', all_text_lower))

        instructions_count = sum([has_time_limit, has_materials, has_grading])
        if instructions_count == 3:
            print(f"PASS: Component 2 — Instructions section has time limit, allowed materials, grading policy (0.15 pts)")
            total_score += 0.15
        elif instructions_count >= 2:
            print(f"PARTIAL: Component 2 — {instructions_count}/3 instruction elements found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Only {instructions_count}/3 instruction elements. time_limit={has_time_limit}, materials={has_materials}, grading={has_grading}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Part A - Multiple Choice with 5 questions and a-d options (0.20 points)
    try:
        has_part_a = idx_part_a >= 0

        mcq_count = 0
        options_count = 0
        if has_part_a:
            # Determine end of Part A section
            end_idx = idx_part_b if idx_part_b >= 0 else len(doc.paragraphs)
            part_a_section = get_section_text(doc, idx_part_a, end_idx)

            # Count numbered questions (e.g., "1. Which..." or "1) Which...")
            mcq_questions = re.findall(r'(?:^|\n)\s*\d+[\.\)]\s+\S', part_a_section)
            mcq_count = len(mcq_questions)

            # Count option lines: "a)" or "a." followed by text
            options_count = len(re.findall(r'(?:^|\n)\s*[abcd]\)', part_a_section))

        has_5_mcqs = mcq_count >= 5
        has_options = options_count >= 16  # 4 options x 5 questions = at least 16

        if has_part_a and has_5_mcqs and has_options:
            print(f"PASS: Component 3 — Part A has {mcq_count} MCQs with {options_count} option markers (0.20 pts)")
            total_score += 0.20
        elif has_part_a and has_5_mcqs:
            print(f"PARTIAL: Component 3 — Part A has {mcq_count} MCQs but only {options_count} option markers (0.10 pts)")
            total_score += 0.10
        elif has_part_a:
            print(f"PARTIAL: Component 3 — Part A heading found but only {mcq_count} MCQs (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — Part A not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Part B - Short Answer with 3 questions (0.15 points)
    try:
        has_part_b = idx_part_b >= 0

        sa_count = 0
        if has_part_b:
            end_idx = idx_part_c if idx_part_c >= 0 else len(doc.paragraphs)
            part_b_section = get_section_text(doc, idx_part_b, end_idx)
            sa_questions = re.findall(r'(?:^|\n)\s*\d+[\.\)]\s+\S', part_b_section)
            sa_count = len(sa_questions)

        has_3_sa = sa_count >= 3

        if has_part_b and has_3_sa:
            print(f"PASS: Component 4 — Part B has {sa_count} short answer questions (0.15 pts)")
            total_score += 0.15
        elif has_part_b:
            print(f"PARTIAL: Component 4 — Part B heading found but only {sa_count} SA questions (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 4 — Part B not found. has_part_b={has_part_b}, sa_count={sa_count}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Part C - Essay with 1 question (0.10 points)
    try:
        has_part_c = idx_part_c >= 0

        essay_has_question = False
        if has_part_c:
            end_idx = idx_answer_key if idx_answer_key >= 0 else len(doc.paragraphs)
            part_c_section = get_section_text(doc, idx_part_c, end_idx)
            # Check for a numbered question or substantial content
            essay_has_question = bool(re.search(r'\d+[\.\)]\s+\S', part_c_section)) or len(part_c_section.strip()) > 100

        if has_part_c and essay_has_question:
            print(f"PASS: Component 5 — Part C: Essay section with question found (0.10 pts)")
            total_score += 0.10
        elif has_part_c:
            print(f"PARTIAL: Component 5 — Part C heading found but no clear essay question (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — Part C not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Page breaks before each part (0.15 points)
    # Task: "Add a page break before each part" => at least 3 breaks (before Part A, B, C)
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        page_break_count = 0
        for para in doc.paragraphs:
            if para.paragraph_format.page_break_before:
                page_break_count += 1
            for run in para.runs:
                for br in run.element.findall('.//w:br', ns):
                    btype = br.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', '')
                    if btype == 'page':
                        page_break_count += 1

        if page_break_count >= 3:
            print(f"PASS: Component 6 — Found {page_break_count} page breaks (need >= 3) (0.15 pts)")
            total_score += 0.15
        elif page_break_count >= 2:
            print(f"PARTIAL: Component 6 — Found {page_break_count} page breaks (need >= 3) (0.08 pts)")
            total_score += 0.08
        elif page_break_count >= 1:
            print(f"PARTIAL: Component 6 — Found {page_break_count} page break (need >= 3) (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No page breaks found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Answer key for Part A on separate page (0.10 points)
    try:
        has_answer_key = idx_answer_key >= 0

        answers_found = 0
        answer_key_page_break_detected = 0  # 0=no, 1=yes
        if has_answer_key:
            answer_section = get_section_text(doc, idx_answer_key, len(doc.paragraphs))
            # Count answer entries (e.g., "1. b)" or "1. b")
            answers_found = len(re.findall(r'\d+[\.\:]\s*[a-d]\)', answer_section.lower()))
            if answers_found == 0:
                answers_found = len(re.findall(r'\d+[\.\:]\s*[a-d][\)\.]?\s', answer_section.lower()))

            # Check for page break before the answer key heading
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            ak_para = doc.paragraphs[idx_answer_key]
            if ak_para.paragraph_format.page_break_before:
                answer_key_page_break_detected = 1
            # Check previous paragraph for run-level page break
            if idx_answer_key > 0 and answer_key_page_break_detected == 0:
                prev_para = doc.paragraphs[idx_answer_key - 1]
                for run in prev_para.runs:
                    for br in run.element.findall('.//w:br', ns):
                        btype = br.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', '')
                        if btype == 'page':
                            answer_key_page_break_detected = 1

        if has_answer_key and answers_found >= 5 and answer_key_page_break_detected == 1:
            print(f"PASS: Component 7 — Answer key with {answers_found} answers on separate page (0.10 pts)")
            total_score += 0.10
        elif has_answer_key and answers_found >= 3:
            print(f"PARTIAL: Component 7 — Answer key found with {answers_found} answers, page_break={answer_key_page_break_detected} (0.05 pts)")
            total_score += 0.05
        elif has_answer_key:
            print(f"PARTIAL: Component 7 — Answer key heading found but only {answers_found} answers (0.03 pts)")
            total_score += 0.03
        else:
            print(f"FAIL: Component 7 — Answer key not found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
