"""
Initial Setup: Create an employee roster document with a plain table.
Task ID: writer_tm_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading('Employee Roster', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('Acme Corporation — Human Resources Department')
    doc.add_paragraph('')

    # Define table data: 5 columns x 10 data rows + 1 header row
    headers = ['Name', 'Department', 'Title', 'Phone', 'Email']
    data = [
        ['Sarah Chen', 'Engineering', 'Senior Developer', '(555) 234-8901', 'sarah.chen@acmecorp.com'],
        ['Marcus Johnson', 'Marketing', 'Campaign Manager', '(555) 345-6712', 'marcus.j@acmecorp.com'],
        ['Elena Rodriguez', 'Finance', 'Financial Analyst', '(555) 456-7823', 'elena.r@acmecorp.com'],
        ['David Kim', 'Engineering', 'DevOps Engineer', '(555) 567-8934', 'david.kim@acmecorp.com'],
        ['Priya Patel', 'Human Resources', 'HR Specialist', '(555) 678-9045', 'priya.p@acmecorp.com'],
        ['James Wilson', 'Sales', 'Account Executive', '(555) 789-0156', 'james.w@acmecorp.com'],
        ['Aisha Thompson', 'Marketing', 'Content Strategist', '(555) 890-1267', 'aisha.t@acmecorp.com'],
        ['Robert Garcia', 'Finance', 'Budget Coordinator', '(555) 901-2378', 'robert.g@acmecorp.com'],
        ['Lisa Chang', 'Engineering', 'QA Lead', '(555) 012-3489', 'lisa.chang@acmecorp.com'],
        ['Michael Brown', 'Sales', 'Regional Director', '(555) 123-4590', 'michael.b@acmecorp.com'],
    ]

    # Create the table: plain "Table Grid" style, no special formatting
    table = doc.add_table(rows=1 + len(data), cols=len(headers), style='Table Grid')

    # Set header row
    for col_idx, header_text in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # Fill in data rows
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(value)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

    # Set column widths for readability
    widths = [Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.3), Inches(2.2)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
