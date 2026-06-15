"""
Reward Script: Apply 'Box List Blue' autoformat style to employee roster table
Task ID: writer_tm_005
Domain: libreoffice_writer
Scoring:
  Component 1: Header row blue background (0.35 pts)
  Component 2: Header row white text (0.30 pts)
  Component 3: Alternating row shading pattern (0.35 pts)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_005'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the 'Box List Blue' autoformat style was applied to the table.
    Checks:
      1. Header row has blue background shading (~4472C4)
      2. Header row text is white (~FFFFFF)
      3. Alternating data rows have light blue shading pattern
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.oxml.ns import qn
    from math import sqrt

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Verify table exists
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    def hex_to_rgb(hex_str):
        """Convert hex color string to (R, G, B) tuple."""
        hex_str = hex_str.lstrip('#')
        return (int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

    def color_distance(c1, c2):
        """Euclidean distance between two RGB tuples."""
        return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))

    def get_cell_fill(cell):
        """Get the fill color of a cell as hex string, or None."""
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            sh = tcPr.find(qn('w:shd'))
            if sh is not None:
                fill = sh.get(qn('w:fill'))
                if fill and fill.lower() != 'auto':
                    return fill
        return None

    # Component 1: Header row has blue background (0.35 points)
    # Expected: fill ~4472C4 (blue) on all header cells
    # This should FAIL on initial (no shading) and PASS on golden
    try:
        header_blue_count = 0
        expected_blue = hex_to_rgb('4472C4')
        num_header_cells = len(table.rows[0].cells)

        for ci, cell in enumerate(table.rows[0].cells):
            fill = get_cell_fill(cell)
            if fill:
                try:
                    cell_rgb = hex_to_rgb(fill)
                    dist = color_distance(cell_rgb, expected_blue)
                    if dist < 60:  # tolerance for blue shade
                        header_blue_count += 1
                        print(f"  Header cell {ci}: fill={fill} — blue match (dist={dist:.1f})")
                    else:
                        print(f"  Header cell {ci}: fill={fill} — NOT blue (dist={dist:.1f})")
                except Exception:
                    print(f"  Header cell {ci}: fill={fill} — parse error")
            else:
                print(f"  Header cell {ci}: no fill")

        if header_blue_count == num_header_cells:
            print(f"PASS: Component 1 — All {num_header_cells} header cells have blue background (0.35 pts)")
            total_score += 0.35
        elif header_blue_count > 0:
            partial = 0.35 * (header_blue_count / num_header_cells)
            print(f"PARTIAL: Component 1 — {header_blue_count}/{num_header_cells} header cells blue ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No header cells have blue background")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row text is white (0.30 points)
    # Expected: font color ~FFFFFF (white) on header text runs
    # This should FAIL on initial (no color set) and PASS on golden
    try:
        white_text_count = 0
        header_text_runs = 0
        expected_white = hex_to_rgb('FFFFFF')

        for ci, cell in enumerate(table.rows[0].cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    if not run.text.strip():
                        continue
                    header_text_runs += 1
                    if run.font.color and run.font.color.rgb:
                        run_rgb = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                        dist = color_distance(run_rgb, expected_white)
                        if dist < 50:
                            white_text_count += 1
                        else:
                            print(f"  Header run '{run.text}': color={run.font.color.rgb} — NOT white (dist={dist:.1f})")
                    else:
                        print(f"  Header run '{run.text}': no explicit color set")

        if header_text_runs > 0 and white_text_count == header_text_runs:
            print(f"PASS: Component 2 — All {header_text_runs} header text runs are white (0.30 pts)")
            total_score += 0.30
        elif white_text_count > 0:
            partial = 0.30 * (white_text_count / max(header_text_runs, 1))
            print(f"PARTIAL: Component 2 — {white_text_count}/{header_text_runs} runs white ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No header text runs have white color")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Alternating row shading (0.35 points)
    # Expected: Odd data rows (1, 3, 5, 7, 9) have light blue fill (~D6E4F0)
    #           Even data rows (2, 4, 6, 8, 10) have no fill
    # KEY: Only count odd data rows having light blue fill as the task-introduced change.
    #      Even rows having no fill is a precondition (true before task), so it must NOT score.
    # This should FAIL on initial (no shading on any row) and PASS on golden
    try:
        expected_light_blue = hex_to_rgb('D6E4F0')
        shaded_rows_correct = 0
        shaded_rows_total = 0  # count of odd data rows that SHOULD have shading

        for ri in range(1, len(table.rows)):
            fill = get_cell_fill(table.cell(ri, 0))
            is_odd_row = (ri % 2 == 1)  # rows 1, 3, 5, 7, 9

            if is_odd_row:
                # Should have light blue fill — this is the task change
                shaded_rows_total += 1
                if fill:
                    try:
                        row_rgb = hex_to_rgb(fill)
                        dist = color_distance(row_rgb, expected_light_blue)
                        if dist < 60:
                            shaded_rows_correct += 1
                            print(f"  Data row {ri}: fill={fill} — light blue match (dist={dist:.1f})")
                        else:
                            print(f"  Data row {ri}: fill={fill} — NOT light blue (dist={dist:.1f})")
                    except Exception:
                        print(f"  Data row {ri}: fill={fill} — parse error")
                else:
                    print(f"  Data row {ri}: no fill (expected light blue)")
            else:
                # Even rows: just log, do not score (no fill is precondition)
                if fill is None:
                    print(f"  Data row {ri}: no fill — as expected (precondition, not scored)")
                else:
                    print(f"  Data row {ri}: fill={fill} — unexpected shading on even row")

        if shaded_rows_total > 0 and shaded_rows_correct == shaded_rows_total:
            print(f"PASS: Component 3 — All {shaded_rows_total} odd data rows have light blue shading (0.35 pts)")
            total_score += 0.35
        elif shaded_rows_correct > 0:
            partial = 0.35 * (shaded_rows_correct / max(shaded_rows_total, 1))
            print(f"PARTIAL: Component 3 — {shaded_rows_correct}/{shaded_rows_total} odd rows shaded ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No odd data rows have light blue shading")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
