"""
Initial Setup: White paper document with paragraphs NOT yet styled with Body Text
Task ID: writer_para_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_para_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Heading 1
    doc.add_heading('Blockchain Technology in Supply Chain Management', level=1)

    # Paragraph 2: Heading 2
    doc.add_heading('Introduction', level=2)

    # Paragraph 3: Main content paragraph — Default Paragraph Style (NOT Body Text)
    p3 = doc.add_paragraph(
        'Supply chain management faces persistent challenges in transparency, '
        'traceability, and trust among participants. Blockchain technology offers '
        'a promising solution through its decentralized, immutable ledger system.'
    )
    # Keep as Normal / Default Paragraph Style (no explicit style change needed)

    # Paragraph 4: Heading 2
    doc.add_heading('Current Challenges', level=2)

    # Paragraph 5: Main content paragraph — Default Paragraph Style (NOT Body Text)
    p5 = doc.add_paragraph(
        'Traditional supply chains suffer from information asymmetry, where different '
        'participants maintain separate records that are difficult to reconcile. '
        'This leads to disputes, fraud, and inefficiency.'
    )

    # Paragraph 6: Heading 2
    doc.add_heading('Proposed Solution', level=2)

    # Paragraph 7: Main content paragraph — Default Paragraph Style (NOT Body Text)
    p7 = doc.add_paragraph(
        'Our blockchain-based platform creates a single source of truth accessible '
        'to all authorized participants. Each transaction is cryptographically signed '
        'and permanently recorded.'
    )

    # Paragraph 8: Heading 2
    doc.add_heading('Implementation Roadmap', level=2)

    # Paragraph 9: Main content paragraph — Default Paragraph Style (NOT Body Text)
    p9 = doc.add_paragraph(
        'The implementation follows a three-phase approach: pilot program with select '
        'partners, gradual expansion to tier-1 suppliers, and full ecosystem rollout.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
