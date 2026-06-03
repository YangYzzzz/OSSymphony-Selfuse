"""
Initial Setup: Company Timeline media kit document with running text (unformatted)
Task ID: writer_mktg_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'  # VM path — file goes on Desktop
TASK_ID = 'media_kit_timeline'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title paragraph — plain text, no special formatting
    title_para = doc.add_paragraph('Company Timeline')
    # No bold, no centering, no special size — just plain text

    # All 8 milestones as one running paragraph (pre-task state)
    running_text = (
        '2018: Founded in San Francisco by Marcus Chen and Priya Sharma. '
        '2019: Launched first product, secured $2M seed funding. '
        '2020: Reached 100 customers, raised $12M Series A. '
        '2021: International expansion to UK and Germany. '
        '2022: Product awarded \'Best Enterprise Solution\' at TechCrunch. '
        '2023: $45M Series B, 500 employees. '
        '2024: Launched AI-powered analytics module. '
        '2025: Surpassed 2,000 customers, $67M ARR.'
    )
    body_para = doc.add_paragraph(running_text)
    # Plain paragraph — no bold, no indentation, no color formatting

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
