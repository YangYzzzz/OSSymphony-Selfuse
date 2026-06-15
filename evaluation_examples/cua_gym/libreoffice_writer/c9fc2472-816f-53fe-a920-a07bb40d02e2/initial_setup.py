"""
Initial Setup: Duplicate section headings as closing labels
Task ID: writer_edit_014
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'company_intro'
OUTPUT = f'{WORKDIR}/Desktop/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # === Section 1: Company Overview ===
    doc.add_heading('Company Overview', level=1)

    p1 = doc.add_paragraph(
        'Founded in 2005, our company has grown from a small regional firm into '
        'a globally recognized leader in enterprise software solutions. We serve '
        'clients across 40 countries, delivering innovative products that transform '
        'the way businesses operate and compete in the digital age.'
    )

    p2 = doc.add_paragraph(
        'Our portfolio includes cloud-based platforms, AI-powered analytics tools, '
        'and end-to-end supply chain management systems. With over 3,200 employees '
        'and annual revenue exceeding $850 million, we remain committed to excellence, '
        'integrity, and continuous improvement in everything we do.'
    )

    # === Section 2: Our Mission ===
    doc.add_heading('Our Mission', level=1)

    p3 = doc.add_paragraph(
        'Our mission is to empower organizations of every size to achieve their full '
        'potential through technology. We believe that the right tools, combined with '
        'deep industry expertise, can unlock new opportunities and drive sustainable '
        'growth for our clients and their communities.'
    )

    p4 = doc.add_paragraph(
        'We accomplish this by investing heavily in research and development, '
        'fostering a culture of curiosity and collaboration, and building long-term '
        'partnerships with our customers. Every product we create is designed with '
        'the user experience at its core, ensuring that powerful capabilities remain '
        'accessible and intuitive.'
    )

    # === Section 3: Global Reach ===
    doc.add_heading('Global Reach', level=1)

    p5 = doc.add_paragraph(
        'With regional headquarters in Singapore, Frankfurt, and São Paulo, we '
        'maintain a strong local presence that enables us to understand and respond '
        'to the unique needs of each market. Our dedicated teams of consultants '
        'and engineers work closely with clients to deliver solutions tailored to '
        'local regulatory, cultural, and operational requirements.'
    )

    p6 = doc.add_paragraph(
        'Our global partner network spans more than 120 certified resellers and '
        'systems integrators, extending our reach into emerging markets and '
        'specialized industries. In 2024, we expanded our operations into Southeast '
        'Asia and Sub-Saharan Africa, bringing our total client base to over '
        '18,000 organizations worldwide.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
