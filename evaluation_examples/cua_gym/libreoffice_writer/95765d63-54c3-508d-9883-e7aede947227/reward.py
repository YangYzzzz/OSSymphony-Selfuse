"""
Reward Script: Insert alphabetical index with marked index entries
Task ID: writer_tech_034
Domain: libreoffice_writer
Scoring:
  Component 1 — XE index entries for 5 terms (0.50 pts, 0.10 each)
  Component 2 — INDEX field present in document (0.20 pts)
  Component 3 — "Index" heading at end of document (0.20 pts)
  Component 4 — XE entries placed at first occurrence of each term (0.10 pts)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_034'

REQUIRED_TERMS = ['authentication', 'endpoint', 'payload', 'token', 'webhook']


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # =========================================================
    # Component 1: XE index entries for all 5 terms (0.50 pts)
    # Each term marked as XE entry earns 0.10 pts
    # =========================================================
    try:
        # Collect all XE entries from instrText fields
        xe_terms_found = set()
        for para in doc.paragraphs:
            for instr in para._element.findall('.//w:instrText', ns):
                text = instr.text.strip() if instr.text else ''
                if text.startswith('XE') or ' XE ' in text:
                    # Extract the term between quotes
                    import re
                    match = re.search(r'XE\s+"([^"]+)"', text)
                    if match:
                        xe_terms_found.add(match.group(1).lower())

        for term in REQUIRED_TERMS:
            if term in xe_terms_found:
                print(f"PASS: Component 1 — XE entry found for '{term}' (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 1 — XE entry missing for '{term}'")

        print(f"  XE entries found: {sorted(xe_terms_found)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================
    # Component 2: INDEX field present in document (0.20 pts)
    # The document must contain an INDEX field code
    # =========================================================
    try:
        index_found = False
        for para in doc.paragraphs:
            for instr in para._element.findall('.//w:instrText', ns):
                text = instr.text.strip() if instr.text else ''
                if 'INDEX' in text.upper():
                    index_found = True
                    print(f"  INDEX field code: {text}")
                    break
            if index_found:
                break

        if index_found:
            print(f"PASS: Component 2 — INDEX field found in document (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — No INDEX field found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================
    # Component 3: "Index" heading at end of document (0.20 pts)
    # There should be an "Index" heading near the end
    # =========================================================
    try:
        # Check the last several paragraphs for a heading with "Index" text
        index_heading_found = False
        # Look in last 10 paragraphs for an Index heading
        search_range = doc.paragraphs[-10:] if len(doc.paragraphs) >= 10 else doc.paragraphs
        for para in search_range:
            style_name = para.style.name if para.style else ''
            text = para.text.strip().lower()
            if 'heading' in style_name.lower() and 'index' in text:
                index_heading_found = True
                print(f"  Found heading: style='{style_name}', text='{para.text.strip()}'")
                break

        if index_heading_found:
            print(f"PASS: Component 3 — 'Index' heading found at end of document (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — No 'Index' heading found at end of document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================
    # Component 4: XE entries at first occurrence of terms (0.10 pts)
    # Each XE entry should be in a paragraph containing the term
    # =========================================================
    try:
        # For each term, find which paragraph has the XE entry and check
        # that the paragraph text contains the term
        correct_placement = 0
        for para in doc.paragraphs:
            para_text_lower = para.text.lower()
            for instr in para._element.findall('.//w:instrText', ns):
                text = instr.text.strip() if instr.text else ''
                import re
                match = re.search(r'XE\s+"([^"]+)"', text)
                if match:
                    term = match.group(1).lower()
                    if term in para_text_lower:
                        correct_placement += 1

        if correct_placement >= 5:
            print(f"PASS: Component 4 — All 5 XE entries placed in paragraphs containing the term (0.10 pts)")
            total_score += 0.10
        elif correct_placement >= 3:
            partial = round(0.10 * correct_placement / 5, 2)
            print(f"PARTIAL: Component 4 — {correct_placement}/5 XE entries correctly placed ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {correct_placement}/5 XE entries in correct paragraphs")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
