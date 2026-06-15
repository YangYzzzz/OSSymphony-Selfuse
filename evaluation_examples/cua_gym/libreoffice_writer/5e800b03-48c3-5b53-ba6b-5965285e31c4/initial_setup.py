"""
Initial Setup: Create a thesis document with heading-like text all in Normal style
Task ID: writer_bs_074
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # All paragraphs use 'Normal' style (Default Paragraph Style).
    # The task is for the agent to apply Heading 1-4 styles to the correct paragraphs.

    # --- Chapter 1 ---
    doc.add_paragraph('Chapter 1: Introduction', style='Normal')
    doc.add_paragraph(
        'This thesis investigates the application of reinforcement learning techniques '
        'to autonomous navigation systems in urban environments. The rapid advancement '
        'of sensor technology and computational power has opened new avenues for developing '
        'intelligent transportation solutions that can adapt to complex, dynamic scenarios.'
    )

    doc.add_paragraph('1.1 Motivation', style='Normal')
    doc.add_paragraph(
        'Urban traffic congestion costs the global economy an estimated $1.4 trillion annually. '
        'Traditional rule-based traffic management systems struggle to adapt to the increasing '
        'complexity of modern transportation networks. Machine learning approaches, particularly '
        'reinforcement learning, offer promising alternatives that can learn optimal strategies '
        'from experience rather than relying on hand-crafted heuristics.'
    )

    doc.add_paragraph('1.1.1 Problem', style='Normal')
    doc.add_paragraph(
        'Current autonomous navigation algorithms rely heavily on pre-mapped environments '
        'and struggle with novel scenarios such as construction zones, emergency vehicles, '
        'and unusual pedestrian behavior. The lack of generalization capability represents '
        'a critical barrier to widespread deployment of autonomous systems.'
    )

    doc.add_paragraph('1.1.1.1 Detail', style='Normal')
    doc.add_paragraph(
        'Specifically, the problem manifests in three dimensions: spatial uncertainty arising '
        'from sensor noise and occlusions, temporal uncertainty due to the unpredictable '
        'behavior of other road users, and contextual uncertainty from varying environmental '
        'conditions including weather, lighting, and road surface quality.'
    )

    doc.add_paragraph('1.1.2 Context', style='Normal')
    doc.add_paragraph(
        'The research builds upon recent breakthroughs in deep reinforcement learning, '
        'particularly the success of policy gradient methods in continuous control tasks. '
        'The urban navigation domain provides a rich testbed where safety constraints, '
        'multi-agent interactions, and partial observability create challenging optimization landscapes.'
    )

    doc.add_paragraph('1.2 Goals', style='Normal')
    doc.add_paragraph(
        'This research aims to develop a hierarchical reinforcement learning framework '
        'that decomposes the navigation task into manageable sub-problems. The framework '
        'should demonstrate improved generalization to unseen scenarios while maintaining '
        'real-time performance suitable for deployment on embedded hardware platforms.'
    )

    doc.add_paragraph('1.2.1 Scope', style='Normal')
    doc.add_paragraph(
        'The scope of this work encompasses simulation-based training using the CARLA '
        'autonomous driving simulator, transfer learning experiments to real-world datasets, '
        'and comparative analysis against state-of-the-art baseline methods including '
        'A* search, rapidly-exploring random trees, and model predictive control.'
    )

    # --- Chapter 2 ---
    doc.add_paragraph('Chapter 2: Background', style='Normal')
    doc.add_paragraph(
        'This chapter provides the theoretical foundations necessary for understanding '
        'the proposed approach. We review key concepts from reinforcement learning theory, '
        'autonomous navigation, and urban traffic modeling that inform our methodology.'
    )

    doc.add_paragraph('2.1 Related Work', style='Normal')
    doc.add_paragraph(
        'Silver et al. (2016) demonstrated that deep reinforcement learning could achieve '
        'superhuman performance in the game of Go. Subsequent work by Mnih et al. extended '
        'these techniques to continuous control domains. In the autonomous driving context, '
        'Dosovitskiy et al. (2017) introduced the CARLA simulator which has become a standard '
        'benchmark for evaluating navigation algorithms.'
    )

    doc.add_paragraph('2.1.1 Classical Approaches', style='Normal')
    doc.add_paragraph(
        'Traditional path planning algorithms such as Dijkstra\'s algorithm and A* search '
        'provide optimal solutions for static environments but scale poorly with state space '
        'dimensionality. Sampling-based methods like RRT and PRM offer probabilistic '
        'completeness guarantees but may produce suboptimal paths without post-processing.'
    )

    doc.add_paragraph('2.1.1.1 Graph-Based Methods', style='Normal')
    doc.add_paragraph(
        'Graph-based planning methods discretize the environment into a connectivity graph '
        'where nodes represent feasible configurations and edges represent collision-free '
        'transitions. The resolution of discretization creates a fundamental tradeoff between '
        'computational efficiency and solution quality that limits practical applicability.'
    )

    doc.add_paragraph('2.2 Reinforcement Learning Fundamentals', style='Normal')
    doc.add_paragraph(
        'A reinforcement learning agent interacts with an environment modeled as a Markov '
        'Decision Process (MDP). The agent observes states, takes actions according to a '
        'policy, and receives scalar rewards. The objective is to learn a policy that maximizes '
        'the expected cumulative discounted reward over time.'
    )

    # --- Chapter 3 ---
    doc.add_paragraph('Chapter 3: Methodology', style='Normal')
    doc.add_paragraph(
        'We propose a hierarchical reinforcement learning architecture consisting of a '
        'high-level strategic planner and a low-level motion controller. The strategic planner '
        'selects waypoints and behavioral primitives, while the motion controller generates '
        'smooth trajectories that satisfy kinematic and dynamic constraints.'
    )

    doc.add_paragraph('3.1 System Architecture', style='Normal')
    doc.add_paragraph(
        'The system architecture follows a modular design pattern with clear interfaces '
        'between perception, planning, and control modules. The perception module processes '
        'raw sensor data from cameras, LiDAR, and radar into a unified world representation '
        'suitable for decision-making by the planning module.'
    )

    doc.add_paragraph('3.1.1 Perception Module', style='Normal')
    doc.add_paragraph(
        'The perception module employs a multi-sensor fusion approach combining three '
        'complementary sensing modalities. Camera images provide rich semantic information, '
        'LiDAR point clouds offer precise depth measurements, and radar signals enable '
        'robust velocity estimation of surrounding objects.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
