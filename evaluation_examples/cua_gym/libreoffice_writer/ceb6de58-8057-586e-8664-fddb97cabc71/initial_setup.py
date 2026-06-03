"""
Initial Setup: Create a Writer document with multiple Heading 2 paragraphs
Task ID: writer_fs_042
Domain: libreoffice_writer

The document is a quarterly business report with several Heading 2 sections.
Heading 2 paragraphs do NOT have keep_with_next or 0.5cm space_before set.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_042'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Q4 2025 Operational Review", level=1)
    title.paragraph_format.space_after = Pt(12)

    intro = doc.add_paragraph(
        "This document summarizes the operational performance across all divisions "
        "for the fourth quarter of fiscal year 2025. Each department head has contributed "
        "their section highlighting key achievements, challenges, and projections for "
        "the upcoming quarter."
    )
    intro.paragraph_format.space_after = Pt(8)

    # --- Section 1: Engineering ---
    h2_1 = doc.add_heading("Engineering Division Performance", level=2)
    # Explicitly ensure keep_with_next is NOT set (default None/inherited)
    h2_1.paragraph_format.keep_with_next = None
    h2_1.paragraph_format.space_before = Pt(0)

    doc.add_paragraph(
        "The Engineering division delivered 14 major product releases during Q4, "
        "surpassing the quarterly target of 10. The team expanded from 87 to 103 engineers, "
        "with notable hires in the machine learning and platform infrastructure groups."
    )
    doc.add_paragraph(
        "Key metrics included a 23% reduction in deployment failures, bringing the "
        "success rate to 99.2%. Average sprint velocity increased by 18 story points "
        "across all teams. The migration to the new microservices architecture reached "
        "72% completion, ahead of the projected 65% milestone."
    )
    doc.add_paragraph(
        "Notable challenges included supply chain delays for GPU hardware procurement, "
        "which pushed the ML training cluster expansion to January 2026. The team "
        "implemented a temporary cloud-burst strategy to maintain research velocity."
    )

    # --- Section 2: Sales ---
    h2_2 = doc.add_heading("Sales and Revenue Analysis", level=2)
    h2_2.paragraph_format.keep_with_next = None
    h2_2.paragraph_format.space_before = Pt(0)

    doc.add_paragraph(
        "Total revenue for Q4 reached $12.8 million, representing a 15% increase "
        "over Q3 and a 34% year-over-year improvement. Enterprise deals accounted for "
        "67% of new bookings, with an average contract value of $245,000."
    )
    doc.add_paragraph(
        "The sales team closed 52 new accounts during the quarter, with a notable "
        "expansion into the healthcare and financial services verticals. Customer "
        "retention rate held steady at 94.3%, with net revenue retention at 118%."
    )
    doc.add_paragraph(
        "Regional performance showed strong growth in the Asia-Pacific market, which "
        "contributed $3.1 million in new annual recurring revenue. The European market "
        "faced headwinds due to regulatory changes but still exceeded targets by 8%."
    )

    # --- Section 3: Marketing ---
    h2_3 = doc.add_heading("Marketing Campaign Results", level=2)
    h2_3.paragraph_format.keep_with_next = None
    h2_3.paragraph_format.space_before = Pt(0)

    doc.add_paragraph(
        "The marketing team executed 6 major campaigns during Q4, generating over "
        "15,000 marketing qualified leads. The annual user conference attracted "
        "2,400 attendees, a 40% increase from the previous year."
    )
    doc.add_paragraph(
        "Digital marketing efforts saw significant improvements, with organic search "
        "traffic growing 28% quarter-over-quarter. The content marketing program "
        "published 34 blog posts, 8 whitepapers, and 12 case studies. Social media "
        "engagement increased by 45% across all platforms."
    )
    doc.add_paragraph(
        "Brand awareness metrics showed a 12-point lift in unaided recall among "
        "target personas in the enterprise segment. The product launch campaign "
        "for Version 5.0 achieved a 62% open rate on email communications."
    )

    # --- Section 4: Customer Success ---
    h2_4 = doc.add_heading("Customer Success and Support Metrics", level=2)
    h2_4.paragraph_format.keep_with_next = None
    h2_4.paragraph_format.space_before = Pt(0)

    doc.add_paragraph(
        "The Customer Success team managed 347 active enterprise accounts during Q4. "
        "Average time to first response decreased from 4.2 hours to 2.8 hours, "
        "while customer satisfaction scores improved to 4.6 out of 5.0."
    )
    doc.add_paragraph(
        "Support ticket volume increased by 22% due to the Version 5.0 release, "
        "but resolution time improved by 15% thanks to the new knowledge base "
        "and automated triaging system. The team resolved 89% of tickets within "
        "the first contact."
    )

    # --- Section 5: Finance ---
    h2_5 = doc.add_heading("Financial Summary and Projections", level=2)
    h2_5.paragraph_format.keep_with_next = None
    h2_5.paragraph_format.space_before = Pt(0)

    doc.add_paragraph(
        "Operating expenses for Q4 totaled $9.4 million, with personnel costs "
        "representing 68% of the total. The company achieved positive EBITDA of "
        "$1.2 million for the first time, marking a significant milestone."
    )
    doc.add_paragraph(
        "Cash reserves stood at $28.5 million at quarter end, providing approximately "
        "18 months of runway at current burn rates. Accounts receivable averaged "
        "38 days outstanding, an improvement from 45 days in Q3."
    )
    doc.add_paragraph(
        "For Q1 2026, management projects revenue of $14.5 million based on the "
        "current pipeline of $22.3 million in weighted opportunities. Headcount "
        "is planned to increase by 15 positions across engineering and sales."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
