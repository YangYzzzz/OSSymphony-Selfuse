"""
FINAL REWARD SCRIPT - SUCCESS
Task: As I’m reviewing the draft in LibreOffice Writer, I want the second and third sentences of paragraph 1 to pop out. Could you give those two sentences a #00FF00 highlight so they’re easy to spot while I edit?
Generated: 2025-09-10 14:17:52
Status: success
Model: azure-o3
Total Steps: 5
"""

import os
import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

"""Reward script for verifying that the 2nd and 3rd sentences of the
first paragraph are highlighted #00FF00 (Bright Green) while the
1st sentence is NOT highlighted.

Scoring (progressive):
  • 0.45  – second sentence fully highlighted
  • 0.45  – third  sentence fully highlighted
  • 0.10  – first  sentence NOT highlighted
Partial credit is given if sentences are only partially highlighted or
if the first sentence is only slightly highlighted.
The script prints detailed diagnostics and the final score as
    REWARD: X.X
and returns the score as a float so the evaluation harness can pick it
up programmatically.
"""

# ---------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------

def split_sentences(text: str):
    """Very simple sentence splitter keeping punctuation with the sentence.
    Returns a list of tuples: (start_idx, end_idx, sentence_text)
    """
    pattern = r"[^.!?]*[.!?]"  # greedy up to and incl. punctuation
    matches = list(re.finditer(pattern, text))
    sentences = []
    last_end = 0
    for m in matches:
        start, end = m.start(), m.end()
        sentences.append((start, end, m.group(0)))
        last_end = end
    # Any trailing text without terminal punctuation is a sentence, too
    if last_end < len(text):
        sentences.append((last_end, len(text), text[last_end:]))
    return sentences


def sentence_highlight_ratio(run_infos, sent_span):
    """Return proportion of characters in the sentence span that are
    highlighted Bright Green.
    run_infos – list of (run_start, run_end, is_highlighted)
    sent_span – (sent_start, sent_end)
    """
    s_start, s_end = sent_span
    if s_end <= s_start:
        return 0.0
    highlighted_chars = 0
    total_chars = s_end - s_start

    for run_start, run_end, is_highlighted in run_infos:
        # skip non-overlapping runs
        if run_end <= s_start or run_start >= s_end:
            continue
        overlap_start = max(run_start, s_start)
        overlap_end = min(run_end, s_end)
        overlap_len = overlap_end - overlap_start
        if is_highlighted:
            highlighted_chars += overlap_len
    return highlighted_chars / total_chars


# ---------------------------------------------------------------------
# Main verification function
# ---------------------------------------------------------------------

def verify_highlight(file_path: str):
    WEIGHT_SECOND = 0.45
    WEIGHT_THIRD  = 0.45
    WEIGHT_FIRST  = 0.10

    max_score = 1.0
    score = 0.0

    # ---------- File existence & loading (no points) ----------
    if not os.path.exists(file_path):
        print("✗ File not found:", file_path)
        print("REWARD: 0.0")
        return 0.0
    try:
        doc = Document(file_path)
        print(f"✓ Loaded DOCX – paragraphs: {len(doc.paragraphs)}")
    except Exception as e:
        print("✗ Could not load DOCX:", e)
        print("REWARD: 0.0")
        return 0.0

    # ---------- Locate first non-empty paragraph ----------
    para = next((p for p in doc.paragraphs if p.text.strip()), None)
    if para is None:
        print("✗ No non-empty paragraph found")
        print("REWARD: 0.0")
        return 0.0

    para_text = para.text
    print("Paragraph 1 text:", para_text)

    # ---------- Sentence detection ----------
    sentences = split_sentences(para_text)
    if len(sentences) < 3:
        print("✗ Fewer than 3 sentences detected – cannot verify task")
        print("REWARD: 0.0")
        return 0.0

    for idx, (_, _, sent_text) in enumerate(sentences, 1):
        print(f"  Sentence {idx}: \"{sent_text.strip()}\"")

    # ---------- Build run character-range / highlight map ----------
    run_infos = []  # (start_idx, end_idx, is_highlighted)
    pos = 0
    for run in para.runs:
        run_start = pos
        run_end   = pos + len(run.text)
        is_highlighted = run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN
        run_infos.append((run_start, run_end, is_highlighted))
        pos = run_end

    # ---------- Evaluate highlight ratios ----------
    ratios = []
    for idx, (s_start, s_end, _) in enumerate(sentences):
        ratio = sentence_highlight_ratio(run_infos, (s_start, s_end))
        ratios.append(ratio)
        print(f"Highlight ratio sentence {idx+1}: {ratio:.2%}")

    # ---------- Scoring ----------
    # Second sentence (index 1)
    if ratios[1] >= 0.95:
        print("✓ Second sentence fully highlighted Bright Green")
        score += WEIGHT_SECOND
    elif ratios[1] >= 0.5:
        print("△ Second sentence partially highlighted")
        score += WEIGHT_SECOND * 0.6
    else:
        print("✗ Second sentence insufficiently highlighted")

    # Third sentence (index 2)
    if ratios[2] >= 0.95:
        print("✓ Third sentence fully highlighted Bright Green")
        score += WEIGHT_THIRD
    elif ratios[2] >= 0.5:
        print("△ Third sentence partially highlighted")
        score += WEIGHT_THIRD * 0.6
    else:
        print("✗ Third sentence insufficiently highlighted")

    # First sentence (index 0) – should NOT be highlighted
    if ratios[0] <= 0.10:
        print("✓ First sentence not highlighted (as required)")
        score += WEIGHT_FIRST
    elif ratios[0] <= 0.30:
        print("△ First sentence slightly highlighted – minor issue")
        score += WEIGHT_FIRST * 0.3
    else:
        print("✗ First sentence incorrectly highlighted")

    final_score = min(score, max_score)
    print(f"Total score: {final_score} (of {max_score})")
    print(f"REWARD: {final_score}")
    return final_score

# ---------------------------------------------------------------------
# Execute verification on the expected file
# ---------------------------------------------------------------------
if __name__ == "__main__":
    FILE_PATH = "/home/user/as_im_reviewing_the_draft_in_libreoffice_writer_i_want_the_second_and_third_sentences_of_paragraph_1.docx"
    verify_highlight(FILE_PATH)
