"""
Initial Setup: Contract with tracked changes across three sections
Task ID: writer_biz_061
Domain: libreoffice_writer

Creates a contract document with tracked changes (insertions and deletions)
in three sections. The agent must accept changes in sections 1 & 2 and
reject changes in section 3.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_061'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Namespace map for tracked changes
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_MAP = {'w': WORD_NS}


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, bold=False, italic=False, font_name=None, font_size=None):
    """Create a w:r element with optional formatting."""
    r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{text}</w:t></w:r>')
    if bold or italic or font_name or font_size:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        if bold:
            rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        if italic:
            rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
        if font_name:
            rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>'))
        if font_size:
            sz_val = str(int(font_size * 2))  # half-points
            rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{sz_val}"/>'))
        r.insert(0, rPr)
    return r


def add_tracked_paragraph(body, text_before, inserted_text, deleted_text, author, date, rev_id_start):
    """
    Add a paragraph with tracked changes.
    The paragraph will show:
      - text_before (normal)
      - inserted_text (tracked insertion - green underline in UI)
      - deleted_text (tracked deletion - red strikethrough in UI)
    """
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')

    # Normal text before the changes
    if text_before:
        r_normal = make_run_element(text_before)
        p.append(r_normal)

    # Tracked insertion
    if inserted_text:
        ins = parse_xml(
            f'<w:ins {nsdecls("w")} w:id="{rev_id_start}" '
            f'w:author="{author}" w:date="{date}"></w:ins>'
        )
        r_ins = make_run_element(inserted_text)
        ins.append(r_ins)
        p.append(ins)

    # Tracked deletion
    if deleted_text:
        dele = parse_xml(
            f'<w:del {nsdecls("w")} w:id="{rev_id_start + 1}" '
            f'w:author="{author}" w:date="{date}"></w:del>'
        )
        # Deleted text uses w:delText instead of w:t
        r_del = parse_xml(
            f'<w:r {nsdecls("w")}>'
            f'<w:delText xml:space="preserve">{deleted_text}</w:delText>'
            f'</w:r>'
        )
        dele.append(r_del)
        p.append(dele)

    body.append(p)
    return rev_id_start + 2


def add_heading_paragraph(body, text, level=1):
    """Add a heading paragraph using XML."""
    style_map = {0: "Title", 1: "Heading1", 2: "Heading2"}
    style_name = style_map.get(level, "Heading1")
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    pPr = parse_xml(f'<w:pPr {nsdecls("w")}><w:pStyle w:val="{style_name}"/></w:pPr>')
    p.insert(0, pPr)
    r = make_run_element(text)
    p.append(r)
    body.append(p)


def add_normal_paragraph(body, text, bold=False, alignment=None):
    """Add a normal paragraph via XML."""
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    if alignment:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}><w:jc w:val="{alignment}"/></w:pPr>')
        p.insert(0, pPr)
    r = make_run_element(text, bold=bold)
    p.append(r)
    body.append(p)


def add_empty_paragraph(body):
    """Add an empty paragraph."""
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    body.append(p)


def create_initial():
    doc = Document()

    # Set up basic styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Clear default content
    body = doc.element.body
    # Remove default empty paragraph
    for p in body.findall(qn('w:p')):
        body.remove(p)

    author_main = "Elena Rodriguez"
    author_legal = "David Kim"
    date_main = "2025-11-15T10:30:00Z"
    date_legal = "2025-11-18T14:15:00Z"
    rev_id = 1

    # ===== DOCUMENT TITLE =====
    add_heading_paragraph(body, "PROFESSIONAL SERVICES AGREEMENT", level=0)
    add_empty_paragraph(body)
    add_normal_paragraph(body, "Agreement Number: PSA-2025-0847", bold=True)
    add_normal_paragraph(body, "Effective Date: December 1, 2025")
    add_normal_paragraph(body, "Parties: Meridian Consulting Group LLC and Northfield Industries Corp.")
    add_empty_paragraph(body)

    # ===== SECTION 1: SCOPE OF SERVICES =====
    add_heading_paragraph(body, "Section 1: Scope of Services", level=1)
    add_empty_paragraph(body)

    add_normal_paragraph(body, "1.1 The Consultant shall provide the following professional services to the Client:")
    add_empty_paragraph(body)

    # Tracked change 1.1: insertion of "comprehensive" and deletion of "basic"
    p1 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p1.append(make_run_element("(a) Strategic planning and "))
    ins1 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:ins>')
    ins1.append(make_run_element("comprehensive "))
    p1.append(ins1)
    rev_id += 1
    p1.append(make_run_element("business analysis services, including market research, competitive assessment, and "))
    del1 = parse_xml(f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:del>')
    del1.append(parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">preliminary </w:delText></w:r>'))
    p1.append(del1)
    rev_id += 1
    ins1b = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:ins>')
    ins1b.append(make_run_element("detailed "))
    p1.append(ins1b)
    rev_id += 1
    p1.append(make_run_element("financial forecasting."))
    body.append(p1)
    add_empty_paragraph(body)

    # Tracked change 1.2: addition of new clause about technology
    p2 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p2.append(make_run_element("(b) Technology infrastructure evaluation"))
    ins2 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:ins>')
    ins2.append(make_run_element(" and digital transformation roadmap development"))
    p2.append(ins2)
    rev_id += 1
    p2.append(make_run_element(", with quarterly progress reports submitted to the Client."))
    body.append(p2)
    add_empty_paragraph(body)

    # Tracked change 1.3: deletion of outdated delivery method
    p3 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p3.append(make_run_element("(c) Deliverables shall be provided in electronic format"))
    del2 = parse_xml(f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:del>')
    del2.append(parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve"> and three printed copies</w:delText></w:r>'))
    p3.append(del2)
    rev_id += 1
    p3.append(make_run_element(" via the secure project portal."))
    body.append(p3)
    add_empty_paragraph(body)

    add_normal_paragraph(body, "1.2 The Consultant acknowledges that all work product created under this Agreement shall be subject to the Client's review and approval process.")
    add_empty_paragraph(body)

    # ===== SECTION 2: COMPENSATION AND PAYMENT TERMS =====
    add_heading_paragraph(body, "Section 2: Compensation and Payment Terms", level=1)
    add_empty_paragraph(body)

    add_normal_paragraph(body, "2.1 The Client agrees to compensate the Consultant as follows:")
    add_empty_paragraph(body)

    # Tracked change 2.1: rate change
    p4 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p4.append(make_run_element("(a) Base consulting fee: "))
    del3 = parse_xml(f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:del>')
    del3.append(parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">$175 per hour</w:delText></w:r>'))
    p4.append(del3)
    rev_id += 1
    ins3 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:ins>')
    ins3.append(make_run_element("$195 per hour"))
    p4.append(ins3)
    rev_id += 1
    p4.append(make_run_element(", not to exceed 160 hours per calendar month."))
    body.append(p4)
    add_empty_paragraph(body)

    # Tracked change 2.2: payment terms modification
    p5 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p5.append(make_run_element("(b) Payment shall be due within "))
    del4 = parse_xml(f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:del>')
    del4.append(parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">thirty (30)</w:delText></w:r>'))
    p5.append(del4)
    rev_id += 1
    ins4 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:ins>')
    ins4.append(make_run_element("forty-five (45)"))
    p5.append(ins4)
    rev_id += 1
    p5.append(make_run_element(" calendar days of invoice receipt."))
    body.append(p5)
    add_empty_paragraph(body)

    # Tracked change 2.3: added late payment clause
    p6 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p6.append(make_run_element("(c) Travel and accommodation expenses shall be reimbursed at cost"))
    ins5 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_main}" w:date="{date_main}"></w:ins>')
    ins5.append(make_run_element(", subject to pre-approval for amounts exceeding $500"))
    p6.append(ins5)
    rev_id += 1
    p6.append(make_run_element("."))
    body.append(p6)
    add_empty_paragraph(body)

    add_normal_paragraph(body, "2.2 The Consultant shall submit itemized invoices on the first business day of each calendar month for services rendered in the preceding month.")
    add_empty_paragraph(body)

    # ===== SECTION 3: CONFIDENTIALITY AND NON-DISCLOSURE =====
    add_heading_paragraph(body, "Section 3: Confidentiality and Non-Disclosure", level=1)
    add_empty_paragraph(body)

    add_normal_paragraph(body, "3.1 Both parties acknowledge the sensitive nature of information exchanged under this Agreement.")
    add_empty_paragraph(body)

    # Tracked change 3.1: legal team disagreed - tried to weaken confidentiality period
    p7 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p7.append(make_run_element("(a) Confidential Information shall remain protected for a period of "))
    del5 = parse_xml(f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author_legal}" w:date="{date_legal}"></w:del>')
    del5.append(parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">five (5) years</w:delText></w:r>'))
    p7.append(del5)
    rev_id += 1
    ins6 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_legal}" w:date="{date_legal}"></w:ins>')
    ins6.append(make_run_element("two (2) years"))
    p7.append(ins6)
    rev_id += 1
    p7.append(make_run_element(" following the termination of this Agreement."))
    body.append(p7)
    add_empty_paragraph(body)

    # Tracked change 3.2: legal team disagreed - tried to add broad exception
    p8 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p8.append(make_run_element("(b) Neither party shall disclose Confidential Information to any third party without prior written consent"))
    ins7 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_legal}" w:date="{date_legal}"></w:ins>')
    ins7.append(make_run_element(", except when disclosure is deemed commercially necessary by the receiving party"))
    p8.append(ins7)
    rev_id += 1
    p8.append(make_run_element("."))
    body.append(p8)
    add_empty_paragraph(body)

    # Tracked change 3.3: legal team disagreed - tried to remove remedies clause
    p9 = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    p9.append(make_run_element("(c) "))
    del6 = parse_xml(f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author_legal}" w:date="{date_legal}"></w:del>')
    del6.append(parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">The disclosing party shall be entitled to seek injunctive relief and monetary damages for any breach of this confidentiality obligation.</w:delText></w:r>'))
    p9.append(del6)
    rev_id += 1
    ins8 = parse_xml(f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author_legal}" w:date="{date_legal}"></w:ins>')
    ins8.append(make_run_element("Remedies for breach shall be limited to actual demonstrable damages."))
    p9.append(ins8)
    rev_id += 1
    body.append(p9)
    add_empty_paragraph(body)

    add_normal_paragraph(body, "3.2 The obligations set forth in this section shall survive the termination or expiration of this Agreement.")
    add_empty_paragraph(body)

    # ===== SIGNATURES =====
    add_heading_paragraph(body, "Signatures", level=1)
    add_empty_paragraph(body)
    add_normal_paragraph(body, "For Meridian Consulting Group LLC:")
    add_normal_paragraph(body, "_________________________________")
    add_normal_paragraph(body, "Elena Rodriguez, Managing Partner")
    add_normal_paragraph(body, "Date: _______________")
    add_empty_paragraph(body)
    add_normal_paragraph(body, "For Northfield Industries Corp.:")
    add_normal_paragraph(body, "_________________________________")
    add_normal_paragraph(body, "David Kim, General Counsel")
    add_normal_paragraph(body, "Date: _______________")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
