"""
Reward Script: Formal Business Letter with Date Field, Times New Roman 12pt, Bold Subject
Task ID: writer_rd_001
Domain: libreoffice_writer
Scoring:
  Component 1: Document has content (multiple paragraphs with text) — 0.15
  Component 2: Right-aligned date field at the top paragraph — 0.25
  Component 3: Body text uses Times New Roman 12pt — 0.20
  Component 4: Subject line is bold — 0.15
  Component 5: Paragraph spacing ~0.5 cm after paragraphs — 0.15
  Component 6: Letter structure (sender, recipient, salutation, closing) — 0.10
"""

import os
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_001'

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def has_date_field(para):
    """Check if a paragraph contains a DATE field code."""
    fld_chars = para._element.findall('.//w:fldChar', NS)
    instr_texts = para._element.findall('.//w:instrText', NS)
    for instr in instr_texts:
        if instr.text and 'DATE' in instr.text.upper():
            return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have paragraphs (blank doc = initial state)
    non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty_paras) < 3:
        print(f"FAIL: Document has only {len(non_empty_paras)} non-empty paragraphs (need >= 3)")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Document has substantial content (0.15 points)
    # Initial doc has 0 paragraphs; golden has 27 with multiple text-filled ones
    try:
        if len(non_empty_paras) >= 8:
            print(f"PASS: Component 1 — Document has {len(non_empty_paras)} non-empty paragraphs (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Only {len(non_empty_paras)} non-empty paragraphs, expected >= 8")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Right-aligned date field at the top (0.25 points)
    # The first paragraph should be right-aligned and contain a DATE field code
    try:
        first_para = doc.paragraphs[0]
        is_right = first_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT
        has_field = has_date_field(first_para)

        if is_right and has_field:
            print(f"PASS: Component 2 — First paragraph is right-aligned with DATE field (0.25 pts)")
            total_score += 0.25
        elif is_right and first_para.text.strip():
            # Partial: right-aligned with date-like text but no field code
            # Check if text looks like a date
            text = first_para.text.strip()
            import re
            date_pattern = re.compile(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', re.IGNORECASE)
            if date_pattern.search(text):
                print(f"PASS (partial): Component 2 — Right-aligned date text without field code (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 2 — First para is right-aligned but no date content: {repr(text[:50])}")
        elif has_field:
            print(f"FAIL: Component 2 — Has DATE field but not right-aligned (align={first_para.paragraph_format.alignment})")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — First para not right-aligned (align={first_para.paragraph_format.alignment}) and no DATE field")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body text uses Times New Roman 12pt (0.20 points)
    # Check runs in non-empty paragraphs for font name and size
    try:
        total_runs = 0
        matching_font_name = 0
        matching_font_size = 0

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            for run in para.runs:
                if not run.text.strip():
                    continue
                total_runs += 1
                if run.font.name == 'Times New Roman':
                    matching_font_name += 1
                if run.font.size is not None and abs(run.font.size.pt - 12.0) < 0.5:
                    matching_font_size += 1

        if total_runs == 0:
            print(f"FAIL: Component 3 — No text runs found")
        else:
            name_ratio = matching_font_name / total_runs
            size_ratio = matching_font_size / total_runs

            if name_ratio >= 0.8 and size_ratio >= 0.8:
                print(f"PASS: Component 3 — {matching_font_name}/{total_runs} runs Times New Roman, {matching_font_size}/{total_runs} runs 12pt (0.20 pts)")
                total_score += 0.20
            elif name_ratio >= 0.8 or size_ratio >= 0.8:
                print(f"PARTIAL: Component 3 — font name ratio={name_ratio:.2f}, size ratio={size_ratio:.2f} (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 3 — font name ratio={name_ratio:.2f}, size ratio={size_ratio:.2f}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Subject line is bold (0.15 points)
    # Look for a paragraph that starts with "Re:" or "Subject:" or similar, and check bold
    try:
        subject_found = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and (text.lower().startswith('re:') or text.lower().startswith('subject:')):
                subject_found = True
                # Check if runs are bold
                bold_runs = [r for r in para.runs if r.text.strip() and r.font.bold]
                all_text_runs = [r for r in para.runs if r.text.strip()]
                if all_text_runs and len(bold_runs) == len(all_text_runs):
                    print(f"PASS: Component 4 — Subject line '{text[:50]}' is bold (0.15 pts)")
                    total_score += 0.15
                elif bold_runs:
                    print(f"PARTIAL: Component 4 — Subject line partially bold ({len(bold_runs)}/{len(all_text_runs)} runs) (0.08 pts)")
                    total_score += 0.08
                else:
                    print(f"FAIL: Component 4 — Subject line '{text[:50]}' is not bold")
                break

        if not subject_found:
            # Also check for any bold paragraph that might be a subject line
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                all_text_runs = [r for r in para.runs if r.text.strip()]
                if all_text_runs and all(r.font.bold for r in all_text_runs):
                    # A fully bold paragraph could be the subject line
                    if len(text) > 10 and len(text) < 150:
                        print(f"PASS: Component 4 — Found bold subject line: '{text[:50]}' (0.15 pts)")
                        total_score += 0.15
                        subject_found = True
                        break

            if not subject_found:
                print(f"FAIL: Component 4 — No bold subject line found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Paragraph spacing ~0.5 cm (approx 14.15 pt or ~179705 EMU) after paragraphs (0.15 points)
    # 0.5 cm = 180000 EMU approximately (exact: 5mm * 36000 = 180000)
    try:
        paras_with_spacing = 0
        paras_checked = 0
        target_emu = 180000  # 0.5 cm in EMU
        tolerance = 36000    # ~1mm tolerance

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            paras_checked += 1
            sa = para.paragraph_format.space_after
            if sa is not None:
                # sa is already in EMU
                sa_val = int(sa)
                if abs(sa_val - target_emu) < tolerance:
                    paras_with_spacing += 1

        if paras_checked == 0:
            print(f"FAIL: Component 5 — No paragraphs to check")
        else:
            ratio = paras_with_spacing / paras_checked
            if ratio >= 0.7:
                print(f"PASS: Component 5 — {paras_with_spacing}/{paras_checked} paragraphs have ~0.5cm spacing after (0.15 pts)")
                total_score += 0.15
            elif ratio >= 0.4:
                print(f"PARTIAL: Component 5 — {paras_with_spacing}/{paras_checked} paragraphs have ~0.5cm spacing (0.08 pts)")
                total_score += 0.08
            else:
                print(f"FAIL: Component 5 — Only {paras_with_spacing}/{paras_checked} paragraphs have ~0.5cm spacing")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Letter structure — sender address, recipient address, salutation, closing (0.10 points)
    try:
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        full_lower = full_text.lower()

        has_dear = 'dear ' in full_lower
        has_sincerely = any(w in full_lower for w in ['sincerely', 'regards', 'yours truly', 'respectfully'])
        # Check for address-like content (city, state patterns)
        import re
        has_address = bool(re.search(r'[A-Z][a-z]+,\s*[A-Z]{2}\s+\d{5}', full_text))

        structure_score = sum([has_dear, has_sincerely, has_address])
        if structure_score >= 3:
            print(f"PASS: Component 6 — Letter has salutation, closing, and address (0.10 pts)")
            total_score += 0.10
        elif structure_score >= 2:
            print(f"PARTIAL: Component 6 — Letter has {structure_score}/3 structure elements (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — Letter structure incomplete (dear={has_dear}, closing={has_sincerely}, address={has_address})")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
