"""
Initial Setup: Insert chapter number and page number in header
Task ID: writer_fs_063
Domain: libreoffice_writer

Creates a multi-chapter Writer document with outline numbering on Heading 1.
Header is enabled but empty. Multiple pages per chapter for meaningful pagination.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_outline_numbering(doc):
    """Add outline numbering definition so Heading 1 gets chapter numbers 1, 2, 3..."""
    # Create an abstract numbering definition for outline numbering
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part.numbering_definitions._numbering

    # Create abstractNum for outline numbering
    abstract_num = numbering_elm.makeelement(qn('w:abstractNum'), {
        qn('w:abstractNumId'): '0'
    })
    # Multi-level list type
    multi_level = abstract_num.makeelement(qn('w:multiLevelType'), {
        qn('w:val'): 'multilevel'
    })
    abstract_num.append(multi_level)

    # Level 0 (Heading 1): "1", "2", "3"...
    lvl0 = abstract_num.makeelement(qn('w:lvl'), {qn('w:ilvl'): '0'})
    start0 = lvl0.makeelement(qn('w:start'), {qn('w:val'): '1'})
    fmt0 = lvl0.makeelement(qn('w:numFmt'), {qn('w:val'): 'decimal'})
    lvl_text0 = lvl0.makeelement(qn('w:lvlText'), {qn('w:val'): '%1'})
    lvl_jc0 = lvl0.makeelement(qn('w:lvlJc'), {qn('w:val'): 'left'})
    pstyle0 = lvl0.makeelement(qn('w:pStyle'), {qn('w:val'): 'Heading1'})
    lvl0.extend([start0, fmt0, lvl_text0, lvl_jc0, pstyle0])
    abstract_num.append(lvl0)

    numbering_elm.insert(0, abstract_num)

    # Create a num element referencing the abstractNum
    num = numbering_elm.makeelement(qn('w:num'), {qn('w:numId'): '1'})
    abstract_ref = num.makeelement(qn('w:abstractNumId'), {qn('w:val'): '0'})
    num.append(abstract_ref)
    numbering_elm.append(num)

    return '1'  # numId


def apply_heading_numbering(para, num_id):
    """Apply numbering to a heading paragraph."""
    pPr = para._element.get_or_add_pPr()
    numPr = pPr.makeelement(qn('w:numPr'), {})
    ilvl = numPr.makeelement(qn('w:ilvl'), {qn('w:val'): '0'})
    numId_elm = numPr.makeelement(qn('w:numId'), {qn('w:val'): num_id})
    numPr.extend([ilvl, numId_elm])
    pPr.append(numPr)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Enable header but leave it empty ---
    header = section.header
    header.is_linked_to_previous = False
    # Clear any default content
    for p in header.paragraphs:
        p.text = ""

    # --- Setup outline numbering ---
    num_id = add_outline_numbering(doc)

    # --- Chapter content ---
    chapters = [
        {
            "title": "Introduction to Data Analytics",
            "paragraphs": [
                "Data analytics has become an essential discipline in modern business operations. Organizations across industries rely on data-driven insights to make strategic decisions, optimize processes, and gain competitive advantages in rapidly evolving markets.",
                "The field encompasses a wide range of techniques, from basic descriptive statistics to advanced machine learning algorithms. Understanding these fundamentals is crucial for any professional seeking to leverage data effectively within their organization.",
                "In this document, we explore the key concepts, methodologies, and best practices that form the foundation of contemporary data analytics. Each chapter builds upon the previous one, creating a comprehensive guide for practitioners at all levels.",
                "The importance of data quality cannot be overstated. Before any analysis can begin, practitioners must ensure that the underlying data is accurate, complete, and properly formatted. Data cleaning and preprocessing often consume the majority of time in analytics projects.",
                "Statistical thinking provides the framework for making sense of data. Concepts such as distributions, hypothesis testing, and confidence intervals enable analysts to draw meaningful conclusions while accounting for uncertainty and variability inherent in real-world data.",
            ]
        },
        {
            "title": "Data Collection and Preparation",
            "paragraphs": [
                "Effective data collection begins with a clear understanding of the business questions being addressed. The choice of data sources, collection methods, and sampling strategies directly impacts the quality and reliability of subsequent analyses.",
                "Primary data collection methods include surveys, experiments, and direct observation. Each approach has distinct advantages and limitations. Surveys enable broad coverage but may suffer from response bias. Experiments provide strong causal evidence but are often expensive and time-consuming to conduct.",
                "Secondary data sources offer tremendous value for analytics projects. Government databases, industry reports, social media feeds, and sensor networks generate massive volumes of data that can be repurposed for analytical insights without the cost of primary collection.",
                "Data preparation involves transforming raw data into a format suitable for analysis. This includes handling missing values, removing duplicates, standardizing formats, and engineering new features from existing variables. Automated ETL pipelines streamline this process for recurring analytical workflows.",
                "Quality assurance checkpoints should be established throughout the data pipeline. Validation rules, anomaly detection algorithms, and manual spot checks help ensure that errors are caught early before they propagate through downstream analyses and compromise decision-making.",
                "Modern data lakes and warehouses provide scalable infrastructure for storing and processing large datasets. Technologies such as Apache Spark, Snowflake, and BigQuery enable analysts to work with petabyte-scale data using familiar SQL and Python interfaces.",
                "The ethical dimensions of data collection deserve careful attention. Privacy regulations such as GDPR and CCPA impose strict requirements on how personal data is gathered, stored, and processed. Analysts must be familiar with applicable regulations and organizational policies.",
                "Documentation of data lineage and transformation steps is essential for reproducibility. A well-maintained data catalog helps team members understand what data is available, where it came from, and how it has been processed.",
            ]
        },
        {
            "title": "Exploratory Data Analysis",
            "paragraphs": [
                "Exploratory Data Analysis (EDA) is the critical first step in understanding any dataset. Through visualization and summary statistics, analysts uncover patterns, detect anomalies, and formulate hypotheses that guide subsequent modeling efforts.",
                "Visualization tools such as histograms, scatter plots, and box plots reveal the distributional properties of variables. Patterns that are invisible in raw tables become immediately apparent when data is presented graphically.",
                "Correlation analysis identifies relationships between variables, though analysts must always remember that correlation does not imply causation. Confounding variables and spurious correlations can mislead even experienced practitioners.",
                "Dimensionality reduction techniques like PCA and t-SNE help analysts navigate high-dimensional datasets. By projecting data onto lower-dimensional spaces, these methods reveal clusters and structures that would otherwise remain hidden.",
                "The iterative nature of EDA means that initial findings often lead to new questions and additional analyses. Maintaining a structured notebook of observations and hypotheses helps organize this exploratory process.",
                "Summary statistics including mean, median, standard deviation, and quantiles provide a numerical overview of data characteristics. However, relying solely on summary statistics can be misleading; Anscombe's quartet demonstrates that very different datasets can share identical summary measures.",
                "Geographic visualizations add a spatial dimension to analysis. Heat maps, choropleth maps, and point plots reveal regional patterns and clustering that tabular analysis might miss entirely.",
            ]
        },
        {
            "title": "Statistical Modeling Techniques",
            "paragraphs": [
                "Statistical modeling provides the mathematical framework for understanding relationships within data. From simple linear regression to complex hierarchical models, the choice of technique depends on the nature of the data and the questions being asked.",
                "Linear regression remains one of the most widely used statistical techniques. Its interpretability and well-understood properties make it an excellent starting point for modeling continuous outcomes. The assumptions of linearity, independence, homoscedasticity, and normality must be validated.",
                "Logistic regression extends the regression framework to binary classification problems. Applications range from credit scoring and medical diagnosis to customer churn prediction. The model outputs probabilities that can be calibrated and thresholded according to business requirements.",
                "Time series analysis addresses the unique challenges of temporal data. Techniques such as ARIMA, exponential smoothing, and Prophet enable forecasting of future values based on historical patterns, seasonality, and trend components.",
                "Bayesian methods offer a principled approach to incorporating prior knowledge into statistical models. The posterior distribution combines prior beliefs with observed data, providing a complete characterization of parameter uncertainty.",
                "Model validation through cross-validation, holdout testing, and information criteria ensures that models generalize to new data rather than merely memorizing training examples. Overfitting remains one of the most common pitfalls in statistical modeling.",
                "Ensemble methods combine multiple models to achieve superior predictive performance. Random forests, gradient boosting, and stacking leverage the diversity of individual models to reduce both bias and variance.",
                "Feature selection and regularization techniques help manage model complexity. Methods such as LASSO, ridge regression, and elastic net automatically shrink or eliminate irrelevant features, producing more parsimonious and interpretable models.",
            ]
        },
        {
            "title": "Machine Learning Applications",
            "paragraphs": [
                "Machine learning algorithms automate the process of pattern recognition and prediction. Supervised learning, unsupervised learning, and reinforcement learning each address different categories of problems with distinct methodological approaches.",
                "Deep learning has revolutionized fields such as computer vision, natural language processing, and speech recognition. Neural network architectures including CNNs, RNNs, and Transformers achieve state-of-the-art performance on increasingly complex tasks.",
                "Transfer learning enables practitioners to leverage pre-trained models for new tasks, dramatically reducing the data and computation required. Fine-tuning large language models and vision models has become a standard practice in applied machine learning.",
                "MLOps practices ensure that machine learning models are deployed, monitored, and maintained effectively in production environments. Continuous integration, automated testing, and model versioning are essential components of a mature ML engineering workflow.",
                "Responsible AI encompasses fairness, accountability, transparency, and ethical considerations in machine learning systems. Bias detection, explainability methods, and impact assessments help ensure that AI applications serve all users equitably.",
                "The deployment of machine learning models requires careful consideration of latency, scalability, and cost constraints. Edge computing, model compression, and serving infrastructure all play roles in delivering predictions at the required speed and scale.",
            ]
        },
    ]

    for i, chapter in enumerate(chapters):
        if i > 0:
            # Add page break before each new chapter (except the first)
            doc.add_page_break()

        # Add chapter heading with numbering
        heading = doc.add_heading(chapter["title"], level=1)
        apply_heading_numbering(heading, num_id)

        # Add paragraphs
        for text in chapter["paragraphs"]:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(8)
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Liberation Serif"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
