"""
Initial Setup: Meeting minutes document with 3 paragraphs (no strikethrough)
Task ID: osworld_writer_strikethrough_last_para_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Quarterly Team Meeting Minutes", level=1)

    # --- Paragraph 1: Meeting overview ---
    p1 = doc.add_paragraph()
    p1.add_run(
        "The quarterly team meeting was held on March 15, 2025, at 10:00 AM in Conference Room B. "
        "All department heads were present, including representatives from Engineering, Marketing, Finance, and Operations. "
        "The meeting was chaired by Director Amanda Reyes, who opened with a summary of the previous quarter's achievements. "
        "Attendance was recorded at the start of the session by the administrative coordinator, Linda Park."
    )

    # --- Paragraph 2: Project updates (5 sentences, NO strikethrough) ---
    p2 = doc.add_paragraph()
    p2.add_run(
        "The engineering team presented an update on the Alpha platform migration project, noting that Phase 1 had been completed two weeks ahead of schedule. "
    )
    p2.add_run(
        "Marketing reported a 12% increase in qualified leads attributed to the spring campaign launched in February. "
    )
    p2.add_run(
        "It was proposed that the budget allocation for the Beta testing environment be reduced by 15% to offset rising infrastructure costs. "
    )
    p2.add_run(
        "The finance department confirmed that current expenditures remain within the approved annual budget despite the additional overtime costs incurred in Q1. "
    )
    p2.add_run(
        "Operations raised a concern about the delayed delivery of equipment for the new distribution center, requesting an updated timeline from the procurement team."
    )

    # --- Paragraph 3: Action items and closing ---
    p3 = doc.add_paragraph()
    p3.add_run(
        "Action items were assigned at the close of the meeting to each department head, with deadlines set for the following two weeks. "
        "The HR team was asked to circulate the updated remote work policy document by March 22nd for team review. "
        "Director Reyes thanked all attendees for their contributions and reminded everyone of the next scheduled meeting on April 12, 2025."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
