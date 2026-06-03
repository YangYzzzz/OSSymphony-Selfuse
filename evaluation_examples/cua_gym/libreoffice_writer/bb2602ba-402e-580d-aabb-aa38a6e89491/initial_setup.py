"""
Initial Setup: Create a Writer document with 6 footnotes across 3 pages,
numbered sequentially 1-6 with Arabic numerals and continuous counting.
Task ID: writer_bs_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_footnote(paragraph, footnote_text, doc):
    """Add a footnote to a paragraph using OOXML manipulation."""
    # Get the footnotes part, create if needed
    if not hasattr(doc, '_footnotes_added'):
        doc._footnotes_added = 0

    doc._footnotes_added += 1
    fn_id = doc._footnotes_added

    # Access footnotes part
    footnotes_part = doc.part.element.findall(qn('w:body'))[0]

    # Create footnote reference in the paragraph run
    run = paragraph.add_run()
    rPr = run._element.makeelement(qn('w:rPr'), {})
    rStyle = rPr.makeelement(qn('w:rStyle'), {qn('w:val'): 'FootnoteReference'})
    rPr.append(rStyle)
    run._element.append(rPr)

    footnote_ref = run._element.makeelement(qn('w:footnoteReference'), {qn('w:id'): str(fn_id)})
    run._element.append(footnote_ref)

    return fn_id


def create_initial():
    doc = Document()

    # Set up default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Ensure the footnotes part exists and build footnotes via XML
    # We need to create the footnotes part manually

    # --- Page 1 Content ---
    heading1 = doc.add_heading('Quarterly Revenue Analysis Report', level=1)

    p1 = doc.add_paragraph(
        'The fiscal year 2025 has shown remarkable growth across all divisions of '
        'Meridian Global Technologies. The consolidated revenue reached $4.2 billion, '
        'representing a 15.3% increase over the previous fiscal year.'
    )

    p2 = doc.add_paragraph(
        'The North American market continues to be the primary revenue driver, '
        'contributing 62% of total global revenue. European operations accounted for '
        '23%, while the Asia-Pacific region showed the strongest growth trajectory '
        'at 31% year-over-year improvement.'
    )

    p3 = doc.add_paragraph(
        'Research and development expenditures totaled $780 million, with significant '
        'investments in artificial intelligence platforms and cloud infrastructure. '
        'The company filed 247 new patents during the reporting period, a record high '
        'that underscores the commitment to innovation-driven growth.'
    )

    # --- Page 2 Content ---
    doc.add_page_break()
    heading2 = doc.add_heading('Market Expansion Strategy', level=1)

    p4 = doc.add_paragraph(
        'Strategic acquisitions in Southeast Asia have positioned Meridian for '
        'accelerated growth in emerging markets. The acquisition of TechVista Solutions '
        'in Singapore for $340 million provides immediate access to government '
        'technology contracts across ASEAN member states.'
    )

    p5 = doc.add_paragraph(
        'The Latin American division, established in Q2 2024, has already secured '
        'partnerships with 14 major telecommunications providers. Initial projections '
        'suggest this region could contribute $600 million in annual revenue by 2027, '
        'driven primarily by digital transformation initiatives in Brazil and Mexico.'
    )

    p6 = doc.add_paragraph(
        'Supply chain optimization through the proprietary LogiFlow platform reduced '
        'operational costs by 8.7% across all manufacturing facilities. The platform '
        'processes over 2.3 million data points daily to optimize inventory levels '
        'and distribution routes.'
    )

    # --- Page 3 Content ---
    doc.add_page_break()
    heading3 = doc.add_heading('Financial Outlook and Projections', level=1)

    p7 = doc.add_paragraph(
        'Based on current trajectory and market conditions, the executive leadership '
        'team projects consolidated revenue of $4.85 billion for fiscal year 2026. '
        'This projection accounts for anticipated regulatory changes in the European '
        'Union and ongoing geopolitical considerations in the Asia-Pacific region.'
    )

    p8 = doc.add_paragraph(
        'Capital expenditure plans for the next three years total $2.1 billion, '
        'with priority allocations for data center expansion in Northern Virginia '
        'and Frankfurt. The board of directors has approved a 12% increase in the '
        'quarterly dividend, reflecting confidence in sustained cash flow generation.'
    )

    p9 = doc.add_paragraph(
        'Employee headcount grew to 34,500 globally, with targeted hiring in '
        'cybersecurity, machine learning engineering, and quantum computing research. '
        'The voluntary attrition rate decreased to 7.2%, well below the industry '
        'average of 13.8%, attributed to enhanced compensation packages and flexible '
        'work arrangements implemented in early 2025.'
    )

    # Now we need to add footnotes via direct XML manipulation
    # Save first, then reopen and inject footnotes at XML level
    doc.save(OUTPUT)

    # Re-open and inject footnotes via XML manipulation
    _inject_footnotes()

    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def _inject_footnotes():
    """Inject footnotes into the document via direct XML manipulation."""
    import zipfile
    import shutil
    from lxml import etree

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    CT = 'http://schemas.openxmlformats.org/package/2006/content-types'

    nsmap = {'w': W, 'r': R}

    temp_path = OUTPUT + '.tmp'
    shutil.copy(OUTPUT, temp_path)

    # Read document.xml from the zip
    with zipfile.ZipFile(temp_path, 'r') as zin:
        doc_xml = zin.read('word/document.xml')
        rels_xml = zin.read('word/_rels/document.xml.rels')
        content_types_xml = zin.read('[Content_Types].xml')
        all_names = zin.namelist()

    doc_tree = etree.fromstring(doc_xml)
    rels_tree = etree.fromstring(rels_xml)
    ct_tree = etree.fromstring(content_types_xml)

    # --- Create footnotes.xml ---
    footnotes_texts = [
        'Based on audited financial statements prepared in accordance with IFRS 16 standards, reviewed by Deloitte & Touche LLP.',
        'Revenue figures exclude intercompany transactions and one-time restructuring charges of $45 million.',
        'Patent count includes provisional applications filed with the USPTO, EPO, and WIPO during the calendar year.',
        'Acquisition price reflects enterprise value including assumed debt of $52 million from TechVista Solutions.',
        'Partnership agreements subject to regulatory approval in each respective jurisdiction; three pending as of report date.',
        'Projection assumes stable exchange rates and no significant changes to current trade agreements or tariff structures.',
    ]

    fn_root = etree.Element(qn('w:footnotes'), nsmap={'w': W, 'r': R})

    # Separator footnotes (id 0 and 1 are special)
    for special_id, special_type in [(0, 'separator'), (1, 'continuationSeparator')]:
        fn = etree.SubElement(fn_root, qn('w:footnote'), {qn('w:type'): special_type, qn('w:id'): str(special_id)})
        fp = etree.SubElement(fn, qn('w:p'))
        fpr = etree.SubElement(fp, qn('w:pPr'))
        fsp = etree.SubElement(fpr, qn('w:spacing'), {qn('w:after'): '0', qn('w:line'): '240', qn('w:lineRule'): 'auto'})
        fr = etree.SubElement(fp, qn('w:r'))
        if special_type == 'separator':
            sep = etree.SubElement(fr, qn('w:separator'))
        else:
            sep = etree.SubElement(fr, qn('w:continuationSeparator'))

    # Actual footnotes (ids 1-6... but 1 is taken by continuationSeparator, so use 2-7? No, standard is id=1 for continuation, actual footnotes start at 1)
    # Actually in OOXML: special footnotes are type="separator" id="-1" or "0", and type="continuationSeparator" id="1" or "0"
    # Let me use the standard: separator=0, continuation=1, actual footnotes start at 2
    # Wait, re-checking: In standard OOXML, the separator has id="0" type="separator", continuation has id="1" type="continuationSeparator", and actual footnotes start at id="2" but the footnoteReference in body uses the same id.
    # Let me just use ids 2,3,4,5,6,7 for the 6 footnotes.

    for i, fn_text in enumerate(footnotes_texts):
        fn_id = i + 2  # ids 2 through 7
        fn = etree.SubElement(fn_root, qn('w:footnote'), {qn('w:id'): str(fn_id)})
        fp = etree.SubElement(fn, qn('w:p'))

        # Footnote reference mark in the footnote itself
        fr_ref = etree.SubElement(fp, qn('w:r'))
        fr_rpr = etree.SubElement(fr_ref, qn('w:rPr'))
        fr_rstyle = etree.SubElement(fr_rpr, qn('w:rStyle'), {qn('w:val'): 'FootnoteReference'})
        etree.SubElement(fr_ref, qn('w:footnoteRef'))

        # Space
        fr_space = etree.SubElement(fp, qn('w:r'))
        ft = etree.SubElement(fr_space, qn('w:t'), {'{http://www.w3.org/XML/1998/namespace}space': 'preserve'})
        ft.text = ' '

        # Footnote text
        fr_text = etree.SubElement(fp, qn('w:r'))
        ft2 = etree.SubElement(fr_text, qn('w:t'))
        ft2.text = fn_text

    footnotes_xml = etree.tostring(fn_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # --- Insert footnote references in body paragraphs ---
    body = doc_tree.find(qn('w:body'))
    paragraphs = body.findall(qn('w:p'))

    # We want to put footnotes at specific positions:
    # Page 1: paragraphs 1 and 2 (0-indexed: after heading=0, p1=1, p2=2)
    # Page 2: paragraphs 5 and 6 (after page break para, heading, p4, p5)
    # Page 3: paragraphs 9 and 10 (after page break para, heading, p7, p8)

    # Let's find all text paragraphs (non-empty)
    text_paras = []
    for p in paragraphs:
        # Check if paragraph has text
        texts = p.findall('.//' + qn('w:t'))
        has_text = any(t.text and t.text.strip() for t in texts)
        if has_text:
            text_paras.append(p)

    # Footnote placement: indices into text_paras
    # text_paras: [0]=heading1, [1]=p1, [2]=p2, [3]=p3, [4]=heading2, [5]=p4, [6]=p5, [7]=p6, [8]=heading3, [9]=p7, [10]=p8, [11]=p9
    # Place footnotes on: p1(idx1), p3(idx3), p4(idx5), p6(idx7), p7(idx9), p9(idx11)
    fn_placements = [
        (1, 2),   # p1 -> footnote id 2
        (3, 3),   # p3 -> footnote id 3
        (5, 4),   # p4 -> footnote id 4
        (7, 5),   # p6 -> footnote id 5
        (9, 6),   # p7 -> footnote id 6
        (11, 7),  # p9 -> footnote id 7
    ]

    for para_idx, fn_id in fn_placements:
        if para_idx < len(text_paras):
            para = text_paras[para_idx]
            # Add a new run with footnote reference at the end of the paragraph
            run_elem = etree.SubElement(para, qn('w:r'))
            rpr = etree.SubElement(run_elem, qn('w:rPr'))
            rstyle = etree.SubElement(rpr, qn('w:rStyle'), {qn('w:val'): 'FootnoteReference'})
            # Add superscript
            vertAlign = etree.SubElement(rpr, qn('w:vertAlign'), {qn('w:val'): 'superscript'})
            fnref = etree.SubElement(run_elem, qn('w:footnoteReference'), {qn('w:id'): str(fn_id)})

    # --- Add footnote properties to document (Arabic, continuous = defaults, but be explicit) ---
    body = doc_tree.find(qn('w:body'))
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        sectPr = etree.SubElement(body, qn('w:sectPr'))

    # Default: Arabic numerals (decimal) and continuous counting - these are defaults so we don't need to set them
    # But let's be explicit to ensure correct initial state
    fnPr = sectPr.find(qn('w:footnotePr'))
    if fnPr is None:
        fnPr = etree.SubElement(sectPr, qn('w:footnotePr'))
    # numFmt: decimal = Arabic numerals (this is the default)
    numFmt = fnPr.find(qn('w:numFmt'))
    if numFmt is None:
        numFmt = etree.SubElement(fnPr, qn('w:numFmt'))
    numFmt.set(qn('w:val'), 'decimal')
    # numRestart: continuous (default)
    numRestart = fnPr.find(qn('w:numRestart'))
    if numRestart is None:
        numRestart = etree.SubElement(fnPr, qn('w:numRestart'))
    numRestart.set(qn('w:val'), 'continuous')

    # Serialize updated document.xml
    updated_doc_xml = etree.tostring(doc_tree, xml_declaration=True, encoding='UTF-8', standalone=True)

    # --- Update relationships to include footnotes ---
    # Find max rId
    max_rid = 0
    for rel in rels_tree:
        rid = rel.get('Id', '')
        if rid.startswith('rId'):
            try:
                num = int(rid[3:])
                if num > max_rid:
                    max_rid = num
            except ValueError:
                pass
    new_rid = f'rId{max_rid + 1}'

    # Add footnotes relationship
    new_rel = etree.SubElement(rels_tree, 'Relationship')
    new_rel.set('Id', new_rid)
    new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
    new_rel.set('Target', 'footnotes.xml')

    updated_rels_xml = etree.tostring(rels_tree, xml_declaration=True, encoding='UTF-8', standalone=True)

    # --- Update [Content_Types].xml ---
    # Add Override for footnotes.xml
    ct_override = etree.SubElement(ct_tree, 'Override')
    ct_override.set('PartName', '/word/footnotes.xml')
    ct_override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml')

    updated_ct_xml = etree.tostring(ct_tree, xml_declaration=True, encoding='UTF-8', standalone=True)

    # --- Rebuild the zip ---
    import os
    with zipfile.ZipFile(temp_path, 'r') as zin:
        with zipfile.ZipFile(OUTPUT, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                if item == 'word/document.xml':
                    zout.writestr(item, updated_doc_xml)
                elif item == 'word/_rels/document.xml.rels':
                    zout.writestr(item, updated_rels_xml)
                elif item == '[Content_Types].xml':
                    zout.writestr(item, updated_ct_xml)
                else:
                    zout.writestr(item, zin.read(item))
            # Add footnotes.xml
            zout.writestr('word/footnotes.xml', footnotes_xml)

    os.remove(temp_path)
    print(f'Footnotes injected successfully into {OUTPUT}')


create_initial()
