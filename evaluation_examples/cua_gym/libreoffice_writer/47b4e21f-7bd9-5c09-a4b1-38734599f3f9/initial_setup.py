"""
Initial Setup: Thesis document with default page style throughout
Task ID: writer_fs_087
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_087'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Default style setup (standard margins, no custom page styles) ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ============================================================
    # TITLE PAGE (page 1)
    # ============================================================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    run = title_para.add_run("Machine Learning Approaches for\nPredictive Maintenance in Manufacturing Systems")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(48)
    run = subtitle.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(60)
    run = author.add_run("by\n\nElena Maria Rodriguez")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept.paragraph_format.space_before = Pt(48)
    run = dept.add_run("Department of Industrial Engineering\nStanford University\nMarch 2025")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # ============================================================
    # FRONT MATTER - Abstract (page 2)
    # ============================================================
    doc.add_page_break()

    abstract_heading = doc.add_heading("Abstract", level=1)
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    abstract_text = (
        "This dissertation investigates the application of machine learning techniques "
        "to predictive maintenance in modern manufacturing environments. We develop a "
        "novel framework that combines deep learning models with sensor data analytics "
        "to predict equipment failures before they occur. Our approach leverages "
        "Long Short-Term Memory (LSTM) networks, convolutional neural networks (CNNs), "
        "and ensemble methods to analyze multivariate time-series data collected from "
        "industrial IoT sensors deployed across 47 production lines in three "
        "manufacturing facilities."
    )
    p = doc.add_paragraph(abstract_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    abstract_text2 = (
        "The proposed framework achieves a prediction accuracy of 94.7% for failure "
        "events within a 72-hour window, representing a 23% improvement over existing "
        "state-of-the-art methods. Furthermore, we demonstrate that our approach reduces "
        "unplanned downtime by 31% and maintenance costs by 18% compared to traditional "
        "scheduled maintenance programs. These results were validated through a "
        "12-month deployment across all three facilities."
    )
    p2 = doc.add_paragraph(abstract_text2)
    for run in p2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    abstract_text3 = (
        "Keywords: predictive maintenance, machine learning, deep learning, LSTM, "
        "manufacturing systems, IoT sensors, time-series analysis, industry 4.0"
    )
    p3 = doc.add_paragraph(abstract_text3)
    for run in p3.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.italic = True

    # ============================================================
    # FRONT MATTER - Acknowledgments (page 3)
    # ============================================================
    doc.add_page_break()

    ack_heading = doc.add_heading("Acknowledgments", level=1)
    ack_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ack_paragraphs = [
        "I would like to express my sincere gratitude to my advisor, Professor James "
        "Whitfield, for his unwavering support and guidance throughout this research. "
        "His deep expertise in manufacturing systems and his vision for integrating "
        "artificial intelligence into industrial processes were instrumental in shaping "
        "this work.",

        "I am also deeply grateful to the members of my dissertation committee: "
        "Professor Mei-Ling Chen, Professor David Okonkwo, and Dr. Sarah Petersson. "
        "Their insightful feedback and constructive criticism significantly improved "
        "the quality of this research.",

        "Special thanks to the engineers at Siemens Manufacturing (Munich facility), "
        "Toyota Motor Manufacturing (Georgetown, KY), and Samsung Electronics (Suwon) "
        "for providing access to their production systems and for their patience during "
        "our data collection campaigns.",

        "Finally, I would like to thank my family, especially my parents Carlos and "
        "Maria Rodriguez, and my partner Alex Nakamura, for their encouragement and "
        "understanding during the long years of this doctoral journey."
    ]
    for text in ack_paragraphs:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # ============================================================
    # FRONT MATTER - Table of Contents (pages 4-5)
    # ============================================================
    doc.add_page_break()

    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    toc_entries = [
        ("Abstract", "ii"),
        ("Acknowledgments", "iii"),
        ("Table of Contents", "iv"),
        ("List of Figures", "vi"),
        ("List of Tables", "vii"),
        ("", ""),
        ("Chapter 1: Introduction", "1"),
        ("    1.1 Background and Motivation", "1"),
        ("    1.2 Problem Statement", "4"),
        ("    1.3 Research Objectives", "6"),
        ("    1.4 Contributions", "8"),
        ("    1.5 Thesis Organization", "10"),
        ("", ""),
        ("Chapter 2: Literature Review", "12"),
        ("    2.1 Predictive Maintenance Overview", "12"),
        ("    2.2 Machine Learning in Manufacturing", "16"),
        ("    2.3 Deep Learning for Time-Series Analysis", "21"),
        ("    2.4 IoT Sensor Data Processing", "25"),
        ("    2.5 Gaps in Current Research", "28"),
        ("", ""),
        ("Chapter 3: Methodology", "30"),
        ("    3.1 System Architecture", "30"),
        ("    3.2 Data Collection Framework", "33"),
        ("    3.3 Feature Engineering Pipeline", "36"),
        ("    3.4 Model Design and Training", "39"),
        ("    3.5 Evaluation Metrics", "43"),
        ("", ""),
        ("Chapter 4: Experimental Results", "45"),
        ("    4.1 Dataset Description", "45"),
        ("    4.2 Baseline Comparisons", "48"),
        ("    4.3 LSTM Performance Analysis", "50"),
        ("    4.4 Ensemble Model Results", "53"),
        ("", ""),
        ("Chapter 5: Conclusions and Future Work", "55"),
        ("    5.1 Summary of Contributions", "55"),
        ("    5.2 Limitations", "57"),
        ("    5.3 Future Research Directions", "58"),
        ("", ""),
        ("References", "60"),
    ]

    for entry, page in toc_entries:
        if entry == "":
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        run = p.add_run(entry)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        # Add right-aligned page number with tab
        run2 = p.add_run(f"\t{page}")
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(12)

    # ============================================================
    # MAIN BODY - Chapter 1: Introduction (pages 6+)
    # ============================================================
    doc.add_page_break()

    ch1 = doc.add_heading("Chapter 1: Introduction", level=1)

    sec_11 = doc.add_heading("1.1 Background and Motivation", level=2)

    body_texts_ch1 = [
        "The manufacturing industry stands at a critical juncture in its evolution "
        "toward Industry 4.0. As production systems become increasingly complex and "
        "interconnected, the cost of unplanned equipment failures has escalated "
        "dramatically. According to a 2024 report by the International Society of "
        "Automation, unplanned downtime costs industrial manufacturers an estimated "
        "$50 billion annually worldwide, with individual events costing between "
        "$10,000 and $250,000 per hour depending on the production line.",

        "Traditional maintenance strategies fall into two broad categories: reactive "
        "maintenance, where repairs are performed only after a failure occurs, and "
        "preventive maintenance, where components are replaced on fixed schedules "
        "regardless of their actual condition. Both approaches have significant "
        "drawbacks. Reactive maintenance leads to costly unplanned downtime, while "
        "preventive maintenance often results in unnecessary replacements of components "
        "that still have significant remaining useful life.",

        "Predictive maintenance represents a paradigm shift in how organizations "
        "manage their physical assets. By leveraging real-time sensor data and advanced "
        "analytics, predictive maintenance systems can estimate the remaining useful "
        "life of equipment components and schedule maintenance activities precisely "
        "when they are needed. This approach promises to minimize both unexpected "
        "failures and unnecessary maintenance interventions.",
    ]
    for text in body_texts_ch1:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    sec_12 = doc.add_heading("1.2 Problem Statement", level=2)

    problem_texts = [
        "Despite significant advances in sensor technology and data analytics, several "
        "critical challenges remain in the field of predictive maintenance. First, "
        "manufacturing environments generate enormous volumes of multivariate "
        "time-series data from hundreds of sensors simultaneously. Processing this "
        "data in real-time requires efficient algorithms that can handle high-dimensional "
        "inputs without sacrificing prediction accuracy.",

        "Second, equipment failure patterns in manufacturing are often complex and "
        "non-linear, involving interactions between multiple components and operating "
        "conditions. Traditional statistical methods, such as Weibull analysis and "
        "exponential degradation models, struggle to capture these intricate patterns. "
        "Machine learning approaches offer a promising alternative, but their "
        "application to predictive maintenance remains challenging due to the scarcity "
        "of labeled failure data and the class imbalance inherent in reliability datasets.",
    ]
    for text in problem_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    # ============================================================
    # MAIN BODY - Chapter 2: Literature Review
    # ============================================================
    doc.add_page_break()

    ch2 = doc.add_heading("Chapter 2: Literature Review", level=1)

    sec_21 = doc.add_heading("2.1 Predictive Maintenance Overview", level=2)

    lit_texts = [
        "Predictive maintenance has evolved significantly over the past two decades, "
        "driven by advances in sensor technology, computing power, and machine learning "
        "algorithms. Mobley (2002) provided one of the earliest comprehensive frameworks "
        "for condition-based maintenance, establishing the theoretical foundation for "
        "modern predictive approaches. His work emphasized the importance of vibration "
        "analysis, thermography, and oil analysis as primary diagnostic techniques.",

        "More recently, Carvalho et al. (2019) conducted a systematic literature review "
        "of machine learning applications in predictive maintenance, identifying key "
        "trends and research gaps. Their analysis of 215 papers published between 2010 "
        "and 2019 revealed a growing emphasis on deep learning methods, particularly "
        "for processing raw sensor signals without manual feature extraction.",

        "The integration of IoT sensors into manufacturing systems has accelerated the "
        "adoption of data-driven maintenance strategies. Lee et al. (2014) proposed "
        "a cyber-physical systems architecture for smart manufacturing that includes "
        "predictive maintenance as a core capability. Their framework has been widely "
        "adopted in both academic research and industrial implementations.",
    ]
    for text in lit_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    sec_22 = doc.add_heading("2.2 Machine Learning in Manufacturing", level=2)

    ml_texts = [
        "The application of machine learning to manufacturing processes has expanded "
        "rapidly in recent years. Random forests, support vector machines, and gradient "
        "boosting methods have demonstrated strong performance in fault diagnosis and "
        "quality prediction tasks. Zhang et al. (2017) compared 12 different ML "
        "algorithms for bearing fault detection and found that ensemble methods "
        "consistently outperformed individual classifiers.",

        "Transfer learning has emerged as a promising approach to address the challenge "
        "of limited labeled data in manufacturing applications. Shao et al. (2018) "
        "proposed a domain adaptation framework that transfers knowledge from one "
        "manufacturing process to another, reducing the need for extensive data "
        "collection in new deployment environments.",
    ]
    for text in ml_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    # ============================================================
    # MAIN BODY - Chapter 3: Methodology
    # ============================================================
    doc.add_page_break()

    ch3 = doc.add_heading("Chapter 3: Methodology", level=1)

    sec_31 = doc.add_heading("3.1 System Architecture", level=2)

    method_texts = [
        "Our proposed predictive maintenance framework consists of four interconnected "
        "modules: (1) a data acquisition layer that interfaces with industrial IoT "
        "sensors, (2) a feature engineering pipeline that transforms raw sensor readings "
        "into meaningful features, (3) a model ensemble that combines predictions from "
        "multiple deep learning architectures, and (4) a decision support system that "
        "translates model outputs into actionable maintenance recommendations.",

        "The data acquisition layer supports multiple sensor protocols, including "
        "OPC-UA, MQTT, and Modbus TCP/IP. In our deployment, we monitored 127 sensors "
        "per production line, capturing measurements at frequencies ranging from 1 Hz "
        "for temperature sensors to 10 kHz for vibration accelerometers. The total data "
        "volume across all three facilities exceeded 4.7 TB per month.",

        "The feature engineering pipeline implements both time-domain and "
        "frequency-domain transformations. Time-domain features include rolling "
        "statistics (mean, variance, kurtosis, skewness), peak-to-peak amplitude, "
        "and root mean square (RMS) values computed over configurable windows of "
        "50, 100, 500, and 1000 samples. Frequency-domain features are extracted "
        "using Short-Time Fourier Transform (STFT) and Continuous Wavelet Transform "
        "(CWT) with Morlet wavelets.",
    ]
    for text in method_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    sec_32 = doc.add_heading("3.2 Data Collection Framework", level=2)

    data_texts = [
        "Data collection was conducted across three manufacturing facilities over a "
        "period of 18 months (January 2023 to June 2024). The facilities represent "
        "different manufacturing sectors: automotive assembly (Facility A, Toyota "
        "Georgetown), semiconductor fabrication (Facility B, Samsung Suwon), and "
        "industrial automation equipment (Facility C, Siemens Munich).",

        "Each facility was equipped with a standardized sensor suite that included "
        "triaxial accelerometers (PCB Piezotronics 356A45), infrared temperature "
        "sensors (Optris CT LT), current transformers (Fluke i400s), and acoustic "
        "emission sensors (Physical Acoustics Micro30D). Sensors were installed on "
        "critical rotating machinery including motors, pumps, compressors, and "
        "conveyor systems.",
    ]
    for text in data_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    # ============================================================
    # MAIN BODY - Chapter 4: Experimental Results
    # ============================================================
    doc.add_page_break()

    ch4 = doc.add_heading("Chapter 4: Experimental Results", level=1)

    sec_41 = doc.add_heading("4.1 Dataset Description", level=2)

    results_texts = [
        "The final dataset used for model training and evaluation comprised 2.3 million "
        "sensor reading windows, each representing a 10-second snapshot of all monitored "
        "parameters for a given machine. Among these, 8,472 windows were labeled as "
        "pre-failure events (occurring within the 72-hour prediction horizon), yielding "
        "a class imbalance ratio of approximately 270:1.",

        "To address this severe class imbalance, we employed a combination of Synthetic "
        "Minority Over-sampling Technique (SMOTE) and random under-sampling of the "
        "majority class. After resampling, the training set contained approximately "
        "equal numbers of normal and pre-failure samples. Cross-validation experiments "
        "confirmed that this resampling strategy did not introduce significant bias "
        "in the evaluation metrics.",
    ]
    for text in results_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    sec_42 = doc.add_heading("4.2 Baseline Comparisons", level=2)

    # Add a results table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Model', 'Accuracy (%)', 'Precision (%)', 'Recall (%)']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    data_rows = [
        ['Random Forest', '87.3', '82.1', '79.4'],
        ['SVM (RBF Kernel)', '84.6', '78.9', '76.2'],
        ['XGBoost', '89.1', '85.3', '82.7'],
        ['LSTM (Ours)', '93.2', '91.4', '90.8'],
        ['Ensemble (Ours)', '94.7', '93.2', '92.5'],
    ]
    for i, row_data in enumerate(data_rows, 1):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

    baseline_texts = [
        "",
        "Table 4.1 presents the performance comparison between our proposed models and "
        "baseline methods. The ensemble model, which combines LSTM, CNN, and gradient "
        "boosting predictions through a learned attention mechanism, achieves the highest "
        "overall performance across all metrics.",

        "Notably, the LSTM model alone significantly outperforms all traditional machine "
        "learning baselines, demonstrating the advantage of recurrent architectures for "
        "capturing temporal dependencies in sensor data. The ensemble further improves "
        "upon the LSTM by incorporating complementary features extracted by the CNN "
        "branch (spatial patterns in frequency spectrograms) and the XGBoost branch "
        "(hand-crafted statistical features).",
    ]
    for text in baseline_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    # ============================================================
    # MAIN BODY - Chapter 5: Conclusions
    # ============================================================
    doc.add_page_break()

    ch5 = doc.add_heading("Chapter 5: Conclusions and Future Work", level=1)

    sec_51 = doc.add_heading("5.1 Summary of Contributions", level=2)

    conclusion_texts = [
        "This dissertation has made several contributions to the field of predictive "
        "maintenance in manufacturing systems. First, we developed a comprehensive "
        "framework for real-time equipment health monitoring that integrates multiple "
        "sensor modalities and machine learning architectures. Second, we demonstrated "
        "that ensemble deep learning methods significantly outperform traditional "
        "approaches for failure prediction in complex manufacturing environments.",

        "Third, we validated our framework through an extensive 12-month deployment "
        "across three diverse manufacturing facilities, demonstrating both the "
        "generalizability of our approach and its practical impact on maintenance "
        "operations. The 31% reduction in unplanned downtime and 18% reduction in "
        "maintenance costs represent substantial economic benefits for the "
        "participating organizations.",
    ]
    for text in conclusion_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    sec_52 = doc.add_heading("5.2 Limitations", level=2)

    limitation_texts = [
        "While our results are encouraging, several limitations should be acknowledged. "
        "First, our framework was evaluated primarily on rotating machinery, which "
        "exhibits well-characterized degradation patterns. Extension to other equipment "
        "types, such as electronic components or hydraulic systems, may require "
        "modifications to the feature engineering pipeline.",

        "Second, the 72-hour prediction horizon, while practically useful, may not "
        "be sufficient for maintenance planning in all industrial contexts. Some "
        "organizations require longer prediction windows to coordinate spare parts "
        "procurement and maintenance crew scheduling.",
    ]
    for text in limitation_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    sec_53 = doc.add_heading("5.3 Future Research Directions", level=2)

    future_texts = [
        "Several promising research directions emerge from this work. First, the "
        "integration of physics-informed neural networks could improve model "
        "interpretability and generalization by incorporating domain knowledge about "
        "degradation mechanisms. Second, federated learning approaches could enable "
        "collaborative model training across multiple facilities without sharing "
        "sensitive manufacturing data.",

        "Third, the incorporation of natural language processing for analyzing "
        "maintenance logs and work orders could provide additional context for "
        "failure prediction models. Finally, reinforcement learning could be applied "
        "to optimize maintenance scheduling decisions, taking into account production "
        "constraints, spare parts availability, and workforce allocation.",
    ]
    for text in future_texts:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    # ============================================================
    # References
    # ============================================================
    doc.add_page_break()

    ref_heading = doc.add_heading("References", level=1)

    references = [
        "Carvalho, T.P., Soares, F.A., Vita, R., Francisco, R.P., Basto, J.P., & Alcala, S.G. (2019). A systematic literature review of machine learning methods applied to predictive maintenance. Computers & Industrial Engineering, 137, 106024.",
        "Lee, J., Bagheri, B., & Kao, H.A. (2014). A cyber-physical systems architecture for Industry 4.0-based manufacturing systems. Manufacturing Letters, 3, 18-23.",
        "Mobley, R.K. (2002). An Introduction to Predictive Maintenance (2nd ed.). Butterworth-Heinemann.",
        "Shao, H., Jiang, H., Zhang, H., & Liang, T. (2018). Electric locomotive bearing fault diagnosis using a novel convolutional deep belief network. IEEE Transactions on Industrial Electronics, 65(3), 2727-2736.",
        "Zhang, S., Zhang, S., Wang, B., & Habetler, T.G. (2017). Machine learning and deep learning algorithms for bearing fault diagnostics: A comprehensive review. IEEE Access, 8, 29857-29881.",
        "Zhao, R., Yan, R., Chen, Z., Mao, K., Wang, P., & Gao, R.X. (2019). Deep learning and its applications to machine health monitoring. Mechanical Systems and Signal Processing, 115, 213-237.",
        "Jardine, A.K., Lin, D., & Banjevic, D. (2006). A review on machinery diagnostics and prognostics implementing condition-based maintenance. Mechanical Systems and Signal Processing, 20(7), 1483-1510.",
        "Wang, J., Ma, Y., Zhang, L., Gao, R.X., & Wu, D. (2018). Deep learning for smart manufacturing: Methods and applications. Journal of Manufacturing Systems, 48, 144-156.",
        "Susto, G.A., Schirru, A., Pampuri, S., McLoone, S., & Beghi, A. (2015). Machine learning for predictive maintenance: A multiple classifier approach. IEEE Transactions on Industrial Informatics, 11(3), 812-820.",
        "Ren, L., Sun, Y., Cui, J., & Zhang, L. (2020). Bearing remaining useful life prediction based on deep autoencoder and deep neural networks. Journal of Manufacturing Systems, 48, 71-77.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
