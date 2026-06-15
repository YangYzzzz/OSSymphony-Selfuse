"""
Reward Script: Verify thesis page styles (Title_Page, Front_Matter, Main_Body)
Task ID: writer_fs_087
Domain: libreoffice_writer
Scoring:
  C1 (0.15) - Document has 3 sections (page breaks creating distinct sections)
  C2 (0.15) - Section 0 (Title_Page) has ~5cm top margin
  C3 (0.15) - Section 0 (Title_Page) has no header/footer content and is unlinked
  C4 (0.20) - Section 1 (Front_Matter) footer has PAGE field, centered
  C5 (0.15) - Section 1 (Front_Matter) page number format is lowerRoman
  C6 (0.10) - Section 2 (Main_Body) header has chapter text
  C7 (0.10) - Section 2 (Main_Body) footer has PAGE field with decimal format, start=1
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_087'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def has_page_field(paragraph_element):
    """Check if a paragraph contains a PAGE field code."""
    instr_texts = paragraph_element.findall('.//w:instrText', NS)
    for it in instr_texts:
        if it.text and 'PAGE' in it.text.upper():
            return True
    return False


def get_alignment(paragraph_element):
    """Get paragraph alignment from XML."""
    pPr = paragraph_element.find('w:pPr', NS)
    if pPr is not None:
        jc = pPr.find('w:jc', NS)
        if jc is not None:
            return jc.get(qn('w:val'))
    return None


def get_pgnum_format(sectPr):
    """Get page number format (fmt) and start from section properties."""
    pgNumType = sectPr.find('.//w:pgNumType', NS)
    if pgNumType is not None:
        fmt = pgNumType.get(qn('w:fmt'))
        start = pgNumType.get(qn('w:start'))
        return fmt, start
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = doc.sections
    num_sections = len(sections)

    # Component 1: Document has at least 3 sections (0.15 points)
    # Initial has 1 section; golden has 3. This verifies section breaks were added.
    try:
        if num_sections >= 3:
            print(f"PASS: Component 1 - Document has {num_sections} sections (>= 3) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 - Expected >= 3 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Gate: need at least 3 sections for remaining checks
    if num_sections < 3:
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    sec0 = sections[0]
    sec1 = sections[1]
    sec2 = sections[2]

    # Component 2: Section 0 (Title_Page) top margin ~5cm (0.15 points)
    # 5cm = 1800000 EMU. Initial is 914400 EMU (2.54cm). Allow small tolerance.
    try:
        top_margin_emu = sec0.top_margin
        top_margin_cm = top_margin_emu / 360000.0 if top_margin_emu else 0
        # Allow tolerance: 4.5cm to 5.5cm
        if 4.5 <= top_margin_cm <= 5.5:
            print(f"PASS: Component 2 - Section 0 top margin = {top_margin_cm:.2f} cm (~5cm) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 - Section 0 top margin = {top_margin_cm:.2f} cm, expected ~5cm")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Section 0 (Title_Page) no header/footer content, unlinked (0.15 points)
    # Initial has linked header/footer. Golden has unlinked empty header/footer.
    try:
        header0 = sec0.header
        footer0 = sec0.footer
        header0_empty = all(p.text.strip() == '' for p in header0.paragraphs)
        footer0_empty = all(p.text.strip() == '' for p in footer0.paragraphs)
        # Also check no PAGE field in footer
        footer0_no_page = not any(has_page_field(p._element) for p in footer0.paragraphs)
        header0_unlinked = not header0.is_linked_to_previous
        footer0_unlinked = not footer0.is_linked_to_previous

        if header0_empty and footer0_empty and footer0_no_page and header0_unlinked and footer0_unlinked:
            print(f"PASS: Component 3 - Section 0 has empty unlinked header/footer, no page fields (0.15 pts)")
            total_score += 0.15
        else:
            details = []
            if not header0_empty:
                details.append("header not empty")
            if not footer0_empty:
                details.append("footer not empty")
            if not footer0_no_page:
                details.append("footer has PAGE field")
            if not header0_unlinked:
                details.append("header still linked")
            if not footer0_unlinked:
                details.append("footer still linked")
            print(f"FAIL: Component 3 - Section 0 issues: {', '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Section 1 (Front_Matter) footer has PAGE field, centered (0.20 points)
    # Initial has no PAGE field in any footer. Golden section 1 has centered PAGE field.
    try:
        footer1 = sec1.footer
        footer1_has_page = any(has_page_field(p._element) for p in footer1.paragraphs)
        footer1_centered = any(get_alignment(p._element) == 'center' for p in footer1.paragraphs if has_page_field(p._element))

        if footer1_has_page and footer1_centered:
            print(f"PASS: Component 4 - Section 1 footer has centered PAGE field (0.20 pts)")
            total_score += 0.20
        elif footer1_has_page:
            # Partial: has page field but not centered
            print(f"PARTIAL: Component 4 - Section 1 footer has PAGE field but not centered (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 - Section 1 footer missing PAGE field")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Section 1 (Front_Matter) page number format is lowerRoman (0.15 points)
    # Initial has no pgNumType. Golden section 1 has fmt=lowerRoman.
    try:
        fmt1, start1 = get_pgnum_format(sec1._sectPr)
        if fmt1 and fmt1.lower() in ('lowerroman', 'lower-roman'):
            print(f"PASS: Component 5 - Section 1 pgNumType format = {fmt1} (Roman numerals) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 - Section 1 pgNumType format = {fmt1}, expected lowerRoman")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Section 2 (Main_Body) header has chapter text (0.10 points)
    # Initial has no header content. Golden has "Chapter 1: Introduction" or similar chapter text.
    try:
        header2 = sec2.header
        header2_text = ' '.join(p.text.strip() for p in header2.paragraphs).strip()
        header2_has_chapter = 'chapter' in header2_text.lower()

        if header2_has_chapter and len(header2_text) > 5:
            print(f"PASS: Component 6 - Section 2 header has chapter text: '{header2_text}' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 - Section 2 header text: '{header2_text}', expected chapter name")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    # Component 7: Section 2 (Main_Body) footer has PAGE field with decimal format, start=1 (0.10 points)
    # Initial has no footer PAGE field. Golden section 2 has decimal format starting at 1.
    try:
        footer2 = sec2.footer
        footer2_has_page = any(has_page_field(p._element) for p in footer2.paragraphs)
        fmt2, start2 = get_pgnum_format(sec2._sectPr)
        fmt2_is_decimal = fmt2 is not None and fmt2.lower() == 'decimal'
        start2_is_1 = start2 is not None and str(start2) == '1'

        if footer2_has_page and fmt2_is_decimal and start2_is_1:
            print(f"PASS: Component 7 - Section 2 footer has PAGE field, decimal format, start=1 (0.10 pts)")
            total_score += 0.10
        elif footer2_has_page and fmt2_is_decimal:
            print(f"PARTIAL: Component 7 - Section 2 footer has PAGE field + decimal, but start={start2} (0.05 pts)")
            total_score += 0.05
        elif footer2_has_page:
            print(f"PARTIAL: Component 7 - Section 2 footer has PAGE field but fmt={fmt2}, start={start2} (0.03 pts)")
            total_score += 0.03
        else:
            print(f"FAIL: Component 7 - Section 2 footer missing PAGE field")
    except Exception as e:
        print(f"ERROR: Component 7 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
