"""
Reward Script: Remove yellow highlighting from all text segments
Task ID: writer_frd_024
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): All yellow highlights removed (progressive per-segment credit)
  Component 2 (0.2): Highlights removed AND text content preserved
  Component 3 (0.2): Highlights removed AND bold/italic formatting preserved
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_024'

# Known highlighted segments from initial state with their formatting
# (paragraph_index, text_substring, expected_bold, expected_italic)
INITIAL_HIGHLIGHTED = [
    (4, "12% increase in brand awareness", True, None),
    (5, "6% reduction compared to Q4 2024", None, None),
    (7, "48,000 qualified leads", None, True),
    (8, "14 new content creators", None, None),
    (10, "2.3 million unique visitors", True, None),
    (11, "average session duration improved to 3 minutes 45", None, None),
    (13, "$620,000 (43.7% of total budget)", None, None),
    (14, "4.8x return on ad spend", True, True),
    (16, "three new hires", None, None),
    (18, "15% to capitalize on the strong ROAS", None, None),
    (18, "user-generated content and testimonials", None, True),
]


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have 20 paragraphs (sanity check)
    if len(doc.paragraphs) != 20:
        print(f"PRECONDITION FAIL: Expected 20 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All yellow highlights removed (0.6 points)
    # Count how many of the 11 known highlighted segments now have NO highlight.
    # This FAILS on initial_env (all 11 still highlighted) and PASSES on golden.
    try:
        segments_unhighlighted = 0
        for para_idx, text_sub, _, _ in INITIAL_HIGHLIGHTED:
            if para_idx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_idx]
            for run in para.runs:
                if text_sub in run.text:
                    if run.font.highlight_color is None:
                        segments_unhighlighted += 1
                    else:
                        print(f"  Still highlighted: para {para_idx} '{text_sub[:40]}' color={run.font.highlight_color}")
                    break

        if segments_unhighlighted == 11:
            print(f"PASS: Component 1 - All 11 segments have highlighting removed (0.6 pts)")
            total_score += 0.6
        elif segments_unhighlighted > 0:
            partial = 0.6 * (segments_unhighlighted / 11)
            print(f"PARTIAL: Component 1 - {segments_unhighlighted}/11 segments unhighlighted ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - No highlights were removed (0/11)")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Highlights removed AND text content preserved (0.2 points)
    # Only awards points if highlights were removed AND the text is still there.
    # On initial_env: highlights are NOT removed -> 0 points (correct).
    try:
        text_preserved_and_unhighlighted = 0
        for para_idx, text_sub, _, _ in INITIAL_HIGHLIGHTED:
            if para_idx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_idx]
            found_unhighlighted_with_text = False
            for run in para.runs:
                if text_sub in run.text and run.font.highlight_color is None:
                    found_unhighlighted_with_text = True
                    break
            if not found_unhighlighted_with_text:
                # Fallback: text may be in para.text but runs resplit
                if text_sub in para.text:
                    # Check if ANY run containing part of the text is unhighlighted
                    # But only count if no runs have highlight for this text
                    pass  # Don't count partial matches
            if found_unhighlighted_with_text:
                text_preserved_and_unhighlighted += 1

        if text_preserved_and_unhighlighted == 11:
            print(f"PASS: Component 2 - All 11 segments: highlight removed AND text intact (0.2 pts)")
            total_score += 0.2
        elif text_preserved_and_unhighlighted > 0:
            partial = 0.2 * (text_preserved_and_unhighlighted / 11)
            print(f"PARTIAL: Component 2 - {text_preserved_and_unhighlighted}/11 text+unhighlight OK ({partial:.3f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - No segments have both highlight removed and text preserved")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Highlights removed AND formatting preserved (0.2 points)
    # Only awards points if the highlight was removed AND bold/italic match expected.
    # On initial_env: highlights are NOT removed -> 0 points (correct).
    try:
        format_preserved_and_unhighlighted = 0
        total_checked = 0
        for para_idx, text_sub, exp_bold, exp_italic in INITIAL_HIGHLIGHTED:
            if para_idx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_idx]
            for run in para.runs:
                if text_sub in run.text:
                    # Only count if highlight is removed (the task change)
                    if run.font.highlight_color is None:
                        total_checked += 1
                        bold_ok = (run.font.bold == exp_bold)
                        italic_ok = (run.font.italic == exp_italic)
                        if bold_ok and italic_ok:
                            format_preserved_and_unhighlighted += 1
                        else:
                            print(f"  Format mismatch para {para_idx}: bold={run.font.bold}(exp={exp_bold}), italic={run.font.italic}(exp={exp_italic})")
                    break

        if total_checked == 0:
            print(f"FAIL: Component 3 - No unhighlighted runs found to check formatting")
        elif format_preserved_and_unhighlighted == total_checked:
            print(f"PASS: Component 3 - Formatting preserved on all {total_checked} unhighlighted runs (0.2 pts)")
            total_score += 0.2
        else:
            partial = 0.2 * (format_preserved_and_unhighlighted / total_checked)
            print(f"PARTIAL: Component 3 - {format_preserved_and_unhighlighted}/{total_checked} formatting OK ({partial:.3f} pts)")
            total_score += partial
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'

persist_app_state("libreoffice_writer")

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
