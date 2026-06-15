"""
Initial Setup: Find and remove yellow highlighting from document
Task ID: writer_frd_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    from docx.shared import Inches
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -- Title --
    title = doc.add_heading("Quarterly Performance Review: Marketing Division", level=1)

    # -- Subtitle / meta --
    meta = doc.add_paragraph()
    meta.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    meta.paragraph_format.space_after = Pt(6)
    r = meta.add_run("Prepared by: Elena Vasquez, VP of Marketing")
    r.font.size = Pt(11)
    r.font.italic = True
    meta2 = doc.add_paragraph()
    r2 = meta2.add_run("Review Period: January 2025 - March 2025")
    r2.font.size = Pt(11)
    r2.font.italic = True
    meta2.paragraph_format.space_after = Pt(12)

    # -- Section 1: Executive Summary --
    doc.add_heading("1. Executive Summary", level=2)

    p1 = doc.add_paragraph()
    r = p1.add_run("The marketing division achieved a ")
    r.font.size = Pt(11)
    # HIGHLIGHT 1: "12% increase in brand awareness"
    r_h1 = p1.add_run("12% increase in brand awareness")
    r_h1.font.size = Pt(11)
    r_h1.font.bold = True
    r_h1.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p1.add_run(" during Q1, surpassing our target of 8%. Digital campaigns drove the majority of this growth, with social media engagement rising by 23% across all platforms.")
    r.font.size = Pt(11)

    p2 = doc.add_paragraph()
    r = p2.add_run("Total marketing spend for the quarter was $1.42 million, representing a ")
    r.font.size = Pt(11)
    # HIGHLIGHT 2: "6% reduction compared to Q4 2024"
    r_h2 = p2.add_run("6% reduction compared to Q4 2024")
    r_h2.font.size = Pt(11)
    r_h2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p2.add_run(". This efficiency improvement is attributed to the shift toward programmatic advertising and in-house content production.")
    r.font.size = Pt(11)

    # -- Section 2: Campaign Performance --
    doc.add_heading("2. Campaign Performance", level=2)

    p3 = doc.add_paragraph()
    r = p3.add_run("The Spring Launch campaign generated over ")
    r.font.size = Pt(11)
    # HIGHLIGHT 3: "48,000 qualified leads"
    r_h3 = p3.add_run("48,000 qualified leads")
    r_h3.font.size = Pt(11)
    r_h3.font.italic = True
    r_h3.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p3.add_run(" through a combination of email nurture sequences, paid search, and retargeting display ads. The conversion rate from lead to opportunity stood at 3.7%, up from 2.9% in the previous quarter.")
    r.font.size = Pt(11)

    p4 = doc.add_paragraph()
    r = p4.add_run("Our influencer partnership program expanded to include ")
    r.font.size = Pt(11)
    # HIGHLIGHT 4: "14 new content creators"
    r_h4 = p4.add_run("14 new content creators")
    r_h4.font.size = Pt(11)
    r_h4.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p4.add_run(" across the lifestyle and technology verticals. Early metrics suggest an average engagement rate of 4.2%, which is significantly above the industry benchmark of 2.8%.")
    r.font.size = Pt(11)

    # -- Section 3: Digital Analytics --
    doc.add_heading("3. Digital Analytics", level=2)

    p5 = doc.add_paragraph()
    r = p5.add_run("Website traffic reached ")
    r.font.size = Pt(11)
    # HIGHLIGHT 5: "2.3 million unique visitors"
    r_h5 = p5.add_run("2.3 million unique visitors")
    r_h5.font.size = Pt(11)
    r_h5.font.bold = True
    r_h5.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p5.add_run(" in Q1, with organic search accounting for 41% of total sessions. The blog content strategy contributed to a 17% lift in organic traffic month-over-month.")
    r.font.size = Pt(11)

    p6 = doc.add_paragraph()
    r = p6.add_run("Mobile traffic now accounts for 62% of all visits. The ")
    r.font.size = Pt(11)
    # HIGHLIGHT 6: "average session duration improved to 3 minutes 45 seconds"
    r_h6 = p6.add_run("average session duration improved to 3 minutes 45 seconds")
    r_h6.font.size = Pt(11)
    r_h6.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p6.add_run(", a notable improvement from the 2 minute 50 second average in Q4 2024.")
    r.font.size = Pt(11)

    # -- Section 4: Budget Allocation --
    doc.add_heading("4. Budget Allocation", level=2)

    p7 = doc.add_paragraph()
    r = p7.add_run("Digital advertising consumed the largest share at ")
    r.font.size = Pt(11)
    # HIGHLIGHT 7: "$620,000 (43.7% of total budget)"
    r_h7 = p7.add_run("$620,000 (43.7% of total budget)")
    r_h7.font.size = Pt(11)
    r_h7.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p7.add_run(". Content production and creative services received $310,000, while events and sponsorships were allocated $215,000. The remaining $275,000 covered team salaries, tools, and overhead.")
    r.font.size = Pt(11)

    p8 = doc.add_paragraph()
    r = p8.add_run("The ROI on paid search campaigns reached ")
    r.font.size = Pt(11)
    # HIGHLIGHT 8: "4.8x return on ad spend"
    r_h8 = p8.add_run("4.8x return on ad spend")
    r_h8.font.size = Pt(11)
    r_h8.font.bold = True
    r_h8.font.italic = True
    r_h8.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p8.add_run(", which marks the highest ROAS we have recorded in the past two years.")
    r.font.size = Pt(11)

    # -- Section 5: Team & Personnel --
    doc.add_heading("5. Team & Personnel", level=2)

    p9 = doc.add_paragraph()
    r = p9.add_run("The team expanded with the addition of ")
    r.font.size = Pt(11)
    # HIGHLIGHT 9: "three new hires"
    r_h9 = p9.add_run("three new hires")
    r_h9.font.size = Pt(11)
    r_h9.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p9.add_run(": a Senior Content Strategist (Priya Mehta), a Data Analyst (James O'Brien), and a Social Media Coordinator (Aiko Tanaka). All three have completed onboarding and are now fully integrated into their respective project teams.")
    r.font.size = Pt(11)

    # -- Section 6: Recommendations --
    doc.add_heading("6. Recommendations for Q2 2025", level=2)

    p10 = doc.add_paragraph()
    r = p10.add_run("Based on Q1 results, we recommend increasing the digital advertising budget by ")
    r.font.size = Pt(11)
    # HIGHLIGHT 10: "15% to capitalize on the strong ROAS"
    r_h10 = p10.add_run("15% to capitalize on the strong ROAS")
    r_h10.font.size = Pt(11)
    r_h10.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p10.add_run(". Additionally, we propose launching a customer advocacy program to leverage ")
    r.font.size = Pt(11)
    # HIGHLIGHT 11: "user-generated content and testimonials"
    r_h11 = p10.add_run("user-generated content and testimonials")
    r_h11.font.size = Pt(11)
    r_h11.font.italic = True
    r_h11.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r = p10.add_run(" as a cost-effective growth channel.")
    r.font.size = Pt(11)

    p11 = doc.add_paragraph()
    r = p11.add_run("Finally, the analytics team should prioritize the implementation of an attribution modeling framework to more accurately measure cross-channel campaign effectiveness and inform future budget allocation decisions.")
    r.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
