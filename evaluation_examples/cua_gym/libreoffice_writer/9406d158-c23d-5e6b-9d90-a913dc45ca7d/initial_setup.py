"""
Initial Setup: Create a Writer document with a 6-item numbered list using default indentation.
Task ID: writer_lec_012
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_012'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title paragraph
    title = doc.add_heading("Quarterly Review Action Items", level=1)

    # Add an introductory paragraph
    intro = doc.add_paragraph(
        "The following action items were identified during the Q1 2025 quarterly review meeting "
        "held on March 28, 2025. Each team lead is responsible for completing their assigned tasks "
        "before the next review cycle."
    )
    intro.paragraph_format.space_after = Pt(12)

    # Define 6 numbered list items with realistic content
    items = [
        "Update the customer onboarding documentation to reflect the new portal design changes approved by the UX team on February 14.",
        "Schedule a cross-departmental meeting with Engineering and Marketing to align on the product launch timeline for the Meridian 3.0 release.",
        "Review and finalize the annual budget proposal for the Operations department, incorporating the revised travel policy estimates from Finance.",
        "Migrate the legacy inventory tracking spreadsheets to the new cloud-based asset management system before the April 30 deadline.",
        "Conduct a security audit of all third-party vendor integrations and submit the compliance report to the Information Security Officer.",
        "Prepare a training presentation for the new hires joining in May, covering the standard operating procedures for client communication.",
    ]

    # Create numbered list with default indentation
    # Default indentation: numbers at 0.635 cm (approx 0.25 in), text at 1.27 cm (approx 0.5 in)
    for item_text in items:
        para = doc.add_paragraph(item_text, style='List Number')

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
