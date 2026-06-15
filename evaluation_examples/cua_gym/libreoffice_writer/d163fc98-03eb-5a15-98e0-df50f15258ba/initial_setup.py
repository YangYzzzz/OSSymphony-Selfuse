"""
Initial Setup: Create a 12-page legal brief with headings and plain text references
Task ID: writer_pd_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_filler_paragraphs(doc, count, topic="general"):
    """Add realistic legal filler paragraphs to pad pages."""
    legal_fillers = {
        "general": [
            "The parties hereby acknowledge and agree that the terms set forth in this document shall be binding upon execution by all signatories. Any modifications to the terms herein must be made in writing and signed by all parties involved in this agreement.",
            "Notwithstanding any provision to the contrary contained herein, the obligations of each party under this agreement shall survive the termination or expiration of this agreement to the extent necessary to give effect to such obligations.",
            "In the event of any dispute arising out of or in connection with this agreement, the parties shall first attempt to resolve the matter through good faith negotiation. If such negotiations fail to produce a resolution within thirty (30) calendar days, the matter shall be submitted to binding arbitration in accordance with the rules of the American Arbitration Association.",
            "Each party represents and warrants that it has full power and authority to enter into this agreement and to perform its obligations hereunder. The execution, delivery, and performance of this agreement by each party has been duly authorized by all necessary corporate or organizational action.",
            "This agreement shall be governed by and construed in accordance with the laws of the State of New York, without giving effect to any choice or conflict of law provision or rule that would cause the application of the laws of any other jurisdiction.",
            "The waiver by either party of a breach of any provision of this agreement shall not operate or be construed as a waiver of any subsequent breach. No waiver shall be valid unless in writing and signed by the party granting such waiver.",
            "All notices required or permitted under this agreement shall be in writing and shall be deemed given when delivered personally, sent by certified mail (return receipt requested), or sent by nationally recognized overnight courier to the addresses set forth on the signature page of this agreement.",
            "The prevailing party in any action or proceeding to enforce this agreement shall be entitled to recover its reasonable attorneys' fees, costs, and expenses incurred in connection with such action or proceeding, in addition to any other relief to which such party may be entitled.",
            "Neither party shall be liable for any failure or delay in performing its obligations under this agreement if such failure or delay results from circumstances beyond the reasonable control of that party, including but not limited to acts of God, natural disasters, war, terrorism, riots, embargoes, labor disputes, or governmental actions.",
            "This agreement constitutes the entire understanding between the parties with respect to the subject matter hereof and supersedes all prior negotiations, representations, warranties, commitments, offers, contracts, and writings, whether written or oral, relating to such subject matter.",
        ],
        "liability": [
            "To the maximum extent permitted by applicable law, in no event shall either party be liable to the other party for any indirect, incidental, special, consequential, or punitive damages, including but not limited to loss of profits, loss of data, business interruption, or loss of goodwill, regardless of the theory of liability.",
            "The total aggregate liability of Greenfield Industries under this agreement shall not exceed the total amount of fees actually paid by the Client during the twelve (12) month period immediately preceding the event giving rise to such liability.",
            "The limitations of liability set forth in this section shall apply regardless of whether the alleged liability is based on contract, tort (including negligence), strict liability, or any other legal or equitable theory, and regardless of whether the liable party has been advised of the possibility of such damages.",
            "Notwithstanding the foregoing limitations, nothing in this agreement shall limit either party's liability for (a) death or personal injury caused by negligence, (b) fraud or fraudulent misrepresentation, or (c) any liability which cannot be limited by applicable law.",
        ],
        "confidential": [
            "Each party acknowledges that in the course of performing its obligations under this agreement, it may receive or have access to confidential and proprietary information of the other party. Each party agrees to maintain the confidentiality of such information and to use it solely for the purposes of this agreement.",
            "Confidential Information shall mean all non-public information disclosed by one party to the other, whether orally, in writing, or by inspection of tangible objects, that is designated as confidential or that reasonably should be understood to be confidential given the nature of the information and the circumstances of disclosure.",
            "The obligations of confidentiality set forth herein shall not apply to information that: (a) was already known to the receiving party without restriction; (b) is or becomes publicly available through no fault of the receiving party; (c) is received from a third party without restriction and without breach of any obligation of confidentiality; or (d) is independently developed by the receiving party without use of or reference to the disclosing party's Confidential Information.",
            "Upon termination or expiration of this agreement, each party shall promptly return or destroy all Confidential Information of the other party in its possession, except as required to be retained by applicable law or regulation. A certifying officer of the returning party shall provide written confirmation of such return or destruction upon request.",
        ],
        "ip": [
            "All intellectual property rights in any work product, deliverables, or materials created by Greenfield Industries in the performance of this agreement shall be owned exclusively by the Client upon full payment of all fees due hereunder.",
            "Greenfield Industries hereby assigns to the Client all right, title, and interest in and to any inventions, discoveries, improvements, or works of authorship conceived, developed, or reduced to practice in connection with the services provided under this agreement.",
            "The Client grants to Greenfield Industries a limited, non-exclusive, revocable license to use the Client's trademarks, trade names, and logos solely for the purpose of performing its obligations under this agreement and for no other purpose.",
            "Each party shall indemnify and hold harmless the other party against any claims, damages, losses, or expenses arising from any alleged infringement of intellectual property rights caused by materials provided by the indemnifying party for use under this agreement.",
        ],
        "termination": [
            "This agreement may be terminated by either party upon sixty (60) days prior written notice to the other party. In the event of termination, all outstanding fees for services rendered prior to the effective date of termination shall become immediately due and payable.",
            "Either party may terminate this agreement immediately upon written notice if the other party: (a) materially breaches any provision of this agreement and fails to cure such breach within thirty (30) days after receipt of written notice thereof; (b) becomes insolvent; (c) files a petition in bankruptcy; or (d) makes an assignment for the benefit of creditors.",
            "Upon termination of this agreement for any reason, all licenses granted hereunder shall immediately terminate, and each party shall cease all use of the other party's Confidential Information, intellectual property, and proprietary materials.",
            "The provisions of this agreement that by their nature are intended to survive termination, including but not limited to confidentiality obligations, limitation of liability, indemnification, and dispute resolution, shall survive any termination or expiration of this agreement.",
        ],
    }
    paragraphs = legal_fillers.get(topic, legal_fillers["general"])
    for i in range(count):
        doc.add_paragraph(paragraphs[i % len(paragraphs)])


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ========================================
    # PAGE 1: Title and Introduction
    # ========================================
    title = doc.add_heading('LEGAL BRIEF', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Greenfield Industries, LLC v. Meridian Capital Partners, Inc.')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    case_info = doc.add_paragraph()
    case_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = case_info.add_run('Case No. 2025-CV-04817\nUnited States District Court\nSouthern District of New York')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    # Section 1 heading
    doc.add_heading('Section 1: Introduction and Background', level=1)

    doc.add_paragraph(
        'This legal brief is submitted on behalf of the Plaintiff, Greenfield Industries, LLC '
        '("Greenfield" or "Plaintiff"), in support of its Motion for Summary Judgment against '
        'the Defendant, Meridian Capital Partners, Inc. ("Meridian" or "Defendant"). The facts '
        'of this case are straightforward and demonstrate that there is no genuine dispute of '
        'material fact warranting a trial.'
    )
    doc.add_paragraph(
        'Greenfield Industries is a Delaware limited liability company with its principal place '
        'of business in New York, New York. The company specializes in sustainable energy solutions '
        'and has been operating since March 2018. Meridian Capital Partners is a California '
        'corporation providing financial advisory and investment management services.'
    )

    doc.add_heading('1.1 Procedural History', level=2)
    doc.add_paragraph(
        'On January 15, 2025, Greenfield filed its initial complaint alleging breach of contract, '
        'breach of fiduciary duty, and fraudulent misrepresentation arising from the parties\' '
        'investment management agreement dated June 1, 2023. Meridian filed its Answer and '
        'Counterclaim on March 3, 2025, denying all material allegations and asserting '
        'counterclaims for breach of contract and unjust enrichment.'
    )
    doc.add_paragraph(
        'Discovery commenced on April 1, 2025, and was completed on September 30, 2025. During '
        'the discovery period, over 45,000 documents were produced by the parties collectively, '
        'and twelve depositions were taken, including those of key officers and employees of both '
        'Greenfield and Meridian.'
    )

    # ========================================
    # PAGE 2: Section 2
    # ========================================
    doc.add_heading('Section 2: Statement of Facts', level=1)

    doc.add_paragraph(
        'The following facts are established by the undisputed record evidence, including deposition '
        'testimony, business records, and documentary evidence produced during discovery.'
    )
    doc.add_heading('2.1 The Investment Management Agreement', level=2)
    doc.add_paragraph(
        'On June 1, 2023, Greenfield and Meridian entered into an Investment Management Agreement '
        '(the "Agreement") pursuant to which Meridian agreed to manage a portfolio of investments '
        'totaling approximately $47.5 million on behalf of Greenfield. Under Section 4(a) of the '
        'Agreement, Meridian was required to exercise the standard of care of a reasonably prudent '
        'investment advisor and to act in the best interests of Greenfield at all times.'
    )
    doc.add_heading('2.2 Performance Benchmarks', level=2)
    doc.add_paragraph(
        'Section 7(b) of the Agreement established specific performance benchmarks. Meridian '
        'was required to achieve a minimum annual return of 8.5% on the managed portfolio, net of '
        'all fees and expenses. The Agreement further provided that if the portfolio return fell '
        'below 5.0% in any given quarter, Meridian was obligated to provide a detailed written '
        'explanation and remediation plan to Greenfield within fifteen (15) business days.'
    )
    doc.add_paragraph(
        'During the period from June 2023 through December 2024, the portfolio under Meridian\'s '
        'management generated a total return of negative 12.3%, representing a loss of approximately '
        '$5.85 million. This performance fell dramatically short of the contractually required '
        'benchmarks and was significantly below the relevant market indices during the same period.'
    )

    doc.add_heading('2.3 Unauthorized Transactions', level=2)
    doc.add_paragraph(
        'The record evidence further demonstrates that Meridian conducted a series of unauthorized '
        'transactions involving high-risk derivative instruments, including credit default swaps '
        'and collateralized debt obligations, in direct violation of the investment guidelines '
        'set forth in Exhibit B to the Agreement. These unauthorized transactions accounted for '
        'approximately $18.2 million in portfolio exposure and resulted in realized losses of '
        '$3.7 million.'
    )
    add_filler_paragraphs(doc, 2, "general")

    # ========================================
    # PAGE 3: Contains "see Clause 5.2" plain text
    # ========================================
    doc.add_heading('2.4 Breach Notification and Response', level=2)
    doc.add_paragraph(
        'On January 8, 2025, Greenfield\'s Chief Financial Officer, Rachel Morrison, sent a '
        'formal written notice to Meridian identifying the breaches described above and demanding '
        'immediate corrective action. The notice specifically cited the unauthorized derivative '
        'transactions, the failure to meet performance benchmarks, and the failure to provide '
        'timely quarterly reports as required under the Agreement.'
    )
    doc.add_paragraph(
        'Meridian\'s response, dated January 22, 2025, acknowledged receipt of the notice but '
        'denied any wrongdoing. Meridian asserted that the derivative transactions were within '
        'the scope of its discretionary authority and that market conditions, rather than any '
        'mismanagement, were responsible for the portfolio\'s underperformance.'
    )
    doc.add_paragraph(
        'The fiduciary obligations imposed upon Meridian under the Agreement are substantial '
        'and cannot be discharged merely by pointing to general market conditions. As detailed '
        'in the expert report of Dr. Helena Vasquez (see Clause 5.2 for the detailed analysis '
        'of fiduciary standards), the unauthorized derivative transactions constituted a clear '
        'breach of Meridian\'s duty of care and loyalty to Greenfield.'
    )
    doc.add_paragraph(
        'Furthermore, Meridian failed to disclose material conflicts of interest arising from '
        'its relationships with the counterparties to certain derivative transactions. Internal '
        'emails produced during discovery reveal that Meridian\'s senior portfolio manager, '
        'David Chen, received undisclosed compensation from Apex Financial Group, a counterparty '
        'to three of the unauthorized credit default swap transactions.'
    )
    add_filler_paragraphs(doc, 3, "general")

    # ========================================
    # PAGE 4: Section 3 heading
    # ========================================
    doc.add_heading('Section 3: Legal Standards for Summary Judgment', level=1)

    doc.add_paragraph(
        'Under Federal Rule of Civil Procedure 56(a), summary judgment is appropriate when "the '
        'movant shows that there is no genuine dispute as to any material fact and the movant is '
        'entitled to judgment as a matter of law." Celotex Corp. v. Catrett, 477 U.S. 317, 322 '
        '(1986). The moving party bears the initial burden of demonstrating the absence of a '
        'genuine issue of material fact. Id. at 323.'
    )
    doc.add_heading('3.1 Burden of Proof', level=2)
    doc.add_paragraph(
        'Once the moving party has met its initial burden, the burden shifts to the nonmoving '
        'party to set forth specific facts showing that there is a genuine issue for trial. '
        'Anderson v. Liberty Lobby, Inc., 477 U.S. 242, 248 (1986). The nonmoving party cannot '
        'rest on mere allegations or denials in its pleadings but must present affirmative evidence '
        'from which a jury could reasonably find in its favor. Id. at 256-57.'
    )
    doc.add_heading('3.2 Application to Breach of Contract Claims', level=2)
    doc.add_paragraph(
        'To establish a breach of contract claim under New York law, the plaintiff must demonstrate: '
        '(1) the existence of a valid contract; (2) performance by the plaintiff; (3) breach by '
        'the defendant; and (4) resulting damages. Harris v. Seward Park Housing Corp., 79 A.D.3d '
        '425, 426 (1st Dep\'t 2010). Where the contract terms are clear and unambiguous, '
        'interpretation is a matter of law for the court. Greenfield v. Philles Records, Inc., '
        '98 N.Y.2d 562, 569 (2002).'
    )
    doc.add_heading('3.3 Fiduciary Duty Standards', level=2)
    doc.add_paragraph(
        'Under New York law, an investment manager owes fiduciary duties of care and loyalty to '
        'its clients. These duties require the investment manager to act solely in the best '
        'interests of the client, to exercise reasonable care and skill in managing the client\'s '
        'assets, and to disclose all material conflicts of interest. Birnbaum v. Birnbaum, 73 '
        'N.Y.2d 461, 466 (1989).'
    )
    add_filler_paragraphs(doc, 4, "liability")

    # ========================================
    # PAGE 5: Section 4
    # ========================================
    doc.add_heading('Section 4: Argument', level=1)

    doc.add_paragraph(
        'Greenfield respectfully submits that the undisputed material facts establish each element '
        'of its breach of contract and breach of fiduciary duty claims as a matter of law. '
        'Accordingly, Greenfield is entitled to summary judgment on both claims.'
    )
    doc.add_heading('4.1 Breach of the Investment Management Agreement', level=2)
    doc.add_paragraph(
        'First, there is no dispute that a valid and enforceable contract exists between the '
        'parties. The Investment Management Agreement, executed on June 1, 2023, was signed by '
        'authorized representatives of both Greenfield and Meridian and supported by adequate '
        'consideration, including Greenfield\'s payment of management fees totaling $2.85 million '
        'during the term of the Agreement.'
    )
    doc.add_paragraph(
        'Second, Greenfield performed all of its obligations under the Agreement, including the '
        'timely payment of all management fees and the provision of all required financial '
        'information and documentation to Meridian. There is no allegation or evidence in the '
        'record that Greenfield failed to perform any obligation under the Agreement.'
    )
    doc.add_paragraph(
        'Third, the undisputed evidence establishes that Meridian breached multiple provisions of '
        'the Agreement, including: (a) Section 4(a), by failing to exercise the required standard '
        'of care; (b) Section 5(c), by conducting unauthorized derivative transactions; '
        '(c) Section 7(b), by failing to meet the minimum performance benchmarks; and '
        '(d) Section 9(a), by failing to provide timely quarterly reports.'
    )
    add_filler_paragraphs(doc, 3, "general")

    doc.add_heading('4.2 Damages Analysis', level=2)
    doc.add_paragraph(
        'Greenfield\'s damages expert, Professor William Hartford of Columbia Business School, '
        'has calculated Greenfield\'s total damages at $11,247,000, consisting of: (a) direct '
        'investment losses of $5,850,000; (b) lost opportunity costs of $3,412,000; and '
        '(c) additional costs and expenses of $1,985,000 incurred as a result of Meridian\'s '
        'breaches. Professor Hartford\'s methodology and conclusions are set forth in detail in '
        'his expert report, which has been submitted to the Court.'
    )
    add_filler_paragraphs(doc, 2, "general")

    # ========================================
    # PAGE 6-7: Section 5 with Clause 5.2 heading
    # ========================================
    doc.add_heading('Section 5: Expert Analysis and Testimony', level=1)

    doc.add_heading('5.1 Financial Analysis', level=2)
    doc.add_paragraph(
        'The financial analysis conducted by Professor Hartford demonstrates that the portfolio '
        'losses were directly attributable to Meridian\'s mismanagement rather than general market '
        'conditions. During the relevant period, the S&P 500 index returned 15.2%, the Bloomberg '
        'Aggregate Bond Index returned 4.8%, and a comparable balanced portfolio benchmark returned '
        '9.7%. Meridian\'s managed portfolio, by contrast, lost 12.3%.'
    )
    doc.add_paragraph(
        'Professor Hartford\'s attribution analysis, using standard portfolio decomposition '
        'methodology, demonstrates that approximately 85% of the underperformance relative to the '
        'benchmark was attributable to the unauthorized derivative positions, while the remaining '
        '15% was attributable to excessive portfolio turnover and poor security selection within '
        'the authorized investment categories.'
    )
    add_filler_paragraphs(doc, 4, "ip")

    # Clause 5.2 heading - THIS IS THE TARGET HEADING
    doc.add_heading('Clause 5.2: Fiduciary Standards Analysis', level=2)
    doc.add_paragraph(
        'Dr. Helena Vasquez, a recognized expert in investment management fiduciary standards '
        'with over twenty-five years of experience in the financial services industry, has opined '
        'that Meridian\'s conduct fell below the applicable standard of care in several material '
        'respects. Specifically, Dr. Vasquez identified the following breaches of fiduciary duty:'
    )
    doc.add_paragraph(
        'First, Meridian failed to conduct adequate due diligence on the derivative instruments '
        'it purchased for Greenfield\'s portfolio. The credit default swaps acquired by Meridian '
        'referenced entities with credit ratings below investment grade, a fact that should have '
        'been identified and disclosed to Greenfield prior to execution of the trades.'
    )
    doc.add_paragraph(
        'Second, Meridian\'s senior portfolio manager, David Chen, had a material undisclosed '
        'conflict of interest arising from his receipt of compensation from Apex Financial Group. '
        'This conflict directly affected his investment decisions and constituted a breach of the '
        'duty of loyalty owed to Greenfield.'
    )
    doc.add_paragraph(
        'Third, Meridian failed to implement adequate risk management controls and procedures. '
        'Dr. Vasquez\'s analysis demonstrates that Meridian lacked written policies governing '
        'derivative trading, failed to establish position limits, and did not conduct stress '
        'testing or scenario analysis on the derivative positions, all of which are standard '
        'practices in the investment management industry.'
    )
    add_filler_paragraphs(doc, 3, "confidential")

    # ========================================
    # PAGE 8: Contains "as referenced in Section 3" plain text
    # ========================================
    doc.add_heading('Section 6: Application of Legal Standards to Facts', level=1)

    doc.add_paragraph(
        'Applying the well-established legal standards to the undisputed facts of this case, '
        'Greenfield submits that it is entitled to summary judgment on its breach of contract '
        'and breach of fiduciary duty claims.'
    )
    doc.add_heading('6.1 Contract Breach Is Established as a Matter of Law', level=2)
    doc.add_paragraph(
        'The evidence conclusively demonstrates each element of Greenfield\'s breach of contract '
        'claim. As referenced in Section 3, the legal standard for summary judgment requires '
        'only that Greenfield demonstrate the absence of any genuine dispute of material fact. '
        'The documentary evidence, including the Agreement itself, quarterly account statements, '
        'trade confirmations, and internal communications, establishes that Meridian breached '
        'multiple provisions of the Agreement without any legitimate justification.'
    )
    doc.add_paragraph(
        'Meridian\'s assertion that market conditions excused its performance failures is without '
        'merit. The Agreement contains no force majeure clause or market condition exception that '
        'would relieve Meridian of its obligation to meet the contractually specified performance '
        'benchmarks. Moreover, even if such an exception existed, the evidence demonstrates that '
        'the portfolio losses were primarily attributable to Meridian\'s unauthorized derivative '
        'trading rather than market conditions.'
    )
    add_filler_paragraphs(doc, 3, "liability")

    doc.add_heading('6.2 Fiduciary Breach Analysis', level=2)
    doc.add_paragraph(
        'The undisputed evidence similarly establishes Meridian\'s breach of fiduciary duty. '
        'Meridian\'s duty of care required it to exercise the standard of care of a reasonably '
        'prudent investment advisor. The expert testimony of Dr. Vasquez establishes that Meridian '
        'failed to meet this standard by engaging in unauthorized transactions, failing to conduct '
        'due diligence, and failing to implement adequate risk management controls.'
    )
    add_filler_paragraphs(doc, 2, "confidential")

    # ========================================
    # PAGES 9-10: Section 7
    # ========================================
    doc.add_heading('Section 7: Damages', level=1)

    doc.add_paragraph(
        'Greenfield seeks total damages of $11,247,000, as calculated by its damages expert, '
        'Professor William Hartford. The damages consist of three components: direct investment '
        'losses, lost opportunity costs, and additional costs and expenses.'
    )
    doc.add_heading('7.1 Direct Investment Losses', level=2)
    doc.add_paragraph(
        'The direct investment losses of $5,850,000 represent the actual decline in the portfolio '
        'value attributable to Meridian\'s mismanagement. This figure is derived from a comparison '
        'of the actual portfolio performance with the expected performance under the contractual '
        'benchmarks, after adjusting for market conditions and other relevant factors.'
    )
    doc.add_heading('7.2 Lost Opportunity Costs', level=2)
    doc.add_paragraph(
        'The lost opportunity costs of $3,412,000 represent the investment returns that Greenfield '
        'would have earned had the portfolio been managed in accordance with the terms of the '
        'Agreement and applicable fiduciary standards. Professor Hartford calculated these costs '
        'using a but-for analysis, comparing the actual portfolio returns to the returns of a '
        'hypothetical benchmark portfolio invested in accordance with the Agreement\'s guidelines.'
    )
    doc.add_heading('7.3 Additional Costs and Expenses', level=2)
    doc.add_paragraph(
        'The additional costs and expenses of $1,985,000 include: (a) fees paid to replacement '
        'investment managers ($875,000); (b) forensic accounting and investigation costs '
        '($425,000); (c) legal fees and expenses incurred prior to litigation ($385,000); and '
        '(d) internal costs associated with Greenfield\'s investigation and remediation efforts '
        '($300,000).'
    )
    add_filler_paragraphs(doc, 5, "termination")

    # ========================================
    # PAGES 11-12: Section 8 and Conclusion
    # ========================================
    doc.add_heading('Section 8: Response to Defendant\'s Counterclaims', level=1)

    doc.add_paragraph(
        'Meridian\'s counterclaims for breach of contract and unjust enrichment are without '
        'merit and should be dismissed. Meridian alleges that Greenfield breached the Agreement '
        'by terminating the Agreement prior to the expiration of the initial three-year term. '
        'However, Section 12(b) of the Agreement expressly permits termination upon thirty (30) '
        'days written notice in the event of a material breach by the other party.'
    )
    doc.add_paragraph(
        'As demonstrated above, Meridian materially breached the Agreement in multiple respects. '
        'Greenfield\'s termination notice, dated January 8, 2025, properly invoked Section 12(b) '
        'and identified the specific breaches giving rise to the right of termination. Meridian '
        'was provided with the requisite thirty-day notice period but failed to cure any of the '
        'identified breaches within that period.'
    )
    add_filler_paragraphs(doc, 4, "termination")

    doc.add_heading('8.1 The Unjust Enrichment Claim Fails', level=2)
    doc.add_paragraph(
        'Meridian\'s unjust enrichment counterclaim, which seeks recovery of allegedly unpaid '
        'management fees of $475,000 for the period from October through December 2024, is barred '
        'by the existence of a valid and enforceable contract governing the subject matter. Under '
        'New York law, a claim for unjust enrichment cannot be maintained where a valid contract '
        'governs the same subject matter. Clark-Fitzpatrick, Inc. v. Long Island R.R. Co., 70 '
        'N.Y.2d 382, 388 (1987).'
    )
    add_filler_paragraphs(doc, 3, "general")

    doc.add_heading('Section 9: Conclusion', level=1)

    doc.add_paragraph(
        'For the reasons set forth above, Greenfield Industries, LLC respectfully requests '
        'that this Court grant its Motion for Summary Judgment on its claims for breach of '
        'contract and breach of fiduciary duty, dismiss Meridian Capital Partners, Inc.\'s '
        'counterclaims with prejudice, and award Greenfield damages in the amount of $11,247,000 '
        'plus pre-judgment interest, attorneys\' fees, and costs.'
    )
    doc.add_paragraph()

    respectfully = doc.add_paragraph()
    respectfully.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = respectfully.add_run('Respectfully submitted,')
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    sig = doc.add_paragraph()
    run = sig.add_run('________________________________\n'
                      'Katherine L. Whitfield, Esq.\n'
                      'WHITFIELD & ASSOCIATES LLP\n'
                      '450 Park Avenue, 32nd Floor\n'
                      'New York, NY 10022\n'
                      'Tel: (212) 555-0147\n'
                      'Attorneys for Plaintiff\n'
                      'Greenfield Industries, LLC')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    run = date_para.add_run('Dated: October 15, 2025')
    run.font.name = 'Times New Roman'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
