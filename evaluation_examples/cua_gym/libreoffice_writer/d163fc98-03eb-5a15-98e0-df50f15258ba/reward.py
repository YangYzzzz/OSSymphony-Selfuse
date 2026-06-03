"""
Reward Script: Cross-references in Legal_Brief.docx
Task ID: writer_pd_004
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Bookmark at Clause 5.2 heading target
  Component 2 (0.25): Bookmark at Section 3 heading target
  Component 3 (0.25): Cross-reference field in 'see Clause 5.2' paragraph pointing to Clause 5.2
  Component 4 (0.25): Cross-reference field in 'as referenced in Section 3' paragraph pointing to Section 3
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_004'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify cross-reference task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Identify paragraphs by content
    clause52_heading_idx = None
    section3_heading_idx = None
    see_clause52_idx = None
    as_ref_section3_idx = None

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style_name = p.style.name if p.style else ''

        # Find Clause 5.2 heading
        if 'Clause 5.2' in text and 'Heading' in style_name:
            clause52_heading_idx = i

        # Find Section 3 heading
        if text.startswith('Section 3') and 'Heading' in style_name:
            section3_heading_idx = i

        # Find paragraph containing 'see Clause 5.2'
        if 'see Clause 5.2' in text.lower() or 'see clause 5.2' in text.lower():
            see_clause52_idx = i

        # Find paragraph containing 'as referenced in Section 3'
        if 'as referenced in Section 3' in text or 'as referenced in section 3' in text.lower():
            as_ref_section3_idx = i

    print(f"DEBUG: Clause 5.2 heading at para {clause52_heading_idx}")
    print(f"DEBUG: Section 3 heading at para {section3_heading_idx}")
    print(f"DEBUG: 'see Clause 5.2' text at para {see_clause52_idx}")
    print(f"DEBUG: 'as referenced in Section 3' text at para {as_ref_section3_idx}")

    if clause52_heading_idx is None or section3_heading_idx is None:
        print("FAIL: Could not locate target headings in document")
        print("REWARD: 0.0")
        return 0.0

    if see_clause52_idx is None or as_ref_section3_idx is None:
        print("FAIL: Could not locate cross-reference source paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # Helper: find bookmarks in a paragraph element
    def find_bookmarks(para_element):
        """Return set of bookmark names in this paragraph."""
        bookmarks = set()
        for bm in para_element.findall('.//w:bookmarkStart', ns):
            name = bm.get(f'{{{w_ns}}}name', '')
            if name:
                bookmarks.add(name)
        return bookmarks

    # Helper: find REF field instructions in a paragraph element
    def find_ref_fields(para_element):
        """Return list of REF field target names in this paragraph."""
        refs = []
        instr_texts = para_element.findall('.//w:instrText', ns)
        for it in instr_texts:
            if it.text and 'REF' in it.text:
                # Extract the bookmark name from the instruction
                # Format: " REF _Ref_Clause_5_2 \h "
                parts = it.text.strip().split()
                if len(parts) >= 2 and parts[0] == 'REF':
                    refs.append(parts[1])
        return refs

    # Helper: check if paragraph has fldChar elements (indicating field codes)
    def has_field_codes(para_element):
        fld_chars = para_element.findall('.//w:fldChar', ns)
        return len(fld_chars) > 0

    # Also search all paragraphs for any bookmark that references Clause 5.2 or Section 3
    all_bookmarks = {}  # bookmark_name -> para_index
    for i, p in enumerate(doc.paragraphs):
        for bm_name in find_bookmarks(p._element):
            all_bookmarks[bm_name] = i

    print(f"DEBUG: All bookmarks with '_Ref' or 'Clause' or 'Section': "
          f"{[(k, v) for k, v in all_bookmarks.items() if '_Ref' in k or 'Clause' in k or 'Section' in k or 'clause' in k or 'section' in k]}")

    # Component 1: Bookmark exists at/near Clause 5.2 heading (0.25 points)
    # A bookmark targeting Clause 5.2 should be placed at or near the Clause 5.2 heading paragraph
    try:
        clause52_bookmark_name = None
        # Check for bookmark at the heading paragraph itself
        heading_bookmarks = find_bookmarks(doc.paragraphs[clause52_heading_idx]._element)
        for bm_name in heading_bookmarks:
            # Accept any bookmark that could serve as a cross-ref target for Clause 5.2
            if '_Ref' in bm_name or 'Clause' in bm_name.lower() or 'clause' in bm_name.lower() or '5_2' in bm_name or '52' in bm_name:
                clause52_bookmark_name = bm_name
                break

        # If not found at heading, look for any bookmark at heading para
        if clause52_bookmark_name is None and heading_bookmarks:
            # Accept any non-default bookmark at the heading
            for bm_name in heading_bookmarks:
                if not bm_name.startswith('_GoBack'):
                    clause52_bookmark_name = bm_name
                    break

        if clause52_bookmark_name:
            print(f"PASS: Component 1 - Bookmark '{clause52_bookmark_name}' found at Clause 5.2 heading (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - No bookmark found at Clause 5.2 heading (para {clause52_heading_idx})")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Bookmark exists at/near Section 3 heading (0.25 points)
    try:
        section3_bookmark_name = None
        heading_bookmarks = find_bookmarks(doc.paragraphs[section3_heading_idx]._element)
        for bm_name in heading_bookmarks:
            if '_Ref' in bm_name or 'Section' in bm_name or 'section' in bm_name or '_3' in bm_name:
                section3_bookmark_name = bm_name
                break

        if section3_bookmark_name is None and heading_bookmarks:
            for bm_name in heading_bookmarks:
                if not bm_name.startswith('_GoBack'):
                    section3_bookmark_name = bm_name
                    break

        if section3_bookmark_name:
            print(f"PASS: Component 2 - Bookmark '{section3_bookmark_name}' found at Section 3 heading (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 - No bookmark found at Section 3 heading (para {section3_heading_idx})")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Cross-reference field in 'see Clause 5.2' paragraph (0.25 points)
    # The paragraph should contain a REF field code pointing to the Clause 5.2 bookmark
    try:
        para_elem = doc.paragraphs[see_clause52_idx]._element
        ref_fields = find_ref_fields(para_elem)
        has_fields = has_field_codes(para_elem)

        if ref_fields:
            # Check if any REF field targets a bookmark at the Clause 5.2 heading
            ref_target = ref_fields[0]
            target_para = all_bookmarks.get(ref_target)
            if target_para == clause52_heading_idx:
                print(f"PASS: Component 3 - REF field '{ref_target}' in 'see Clause 5.2' para points to Clause 5.2 heading (0.25 pts)")
                total_score += 0.25
            elif target_para is not None:
                # Points to some bookmark, partial credit if it's close
                print(f"PARTIAL: Component 3 - REF field '{ref_target}' points to para {target_para}, expected {clause52_heading_idx} (0.1 pts)")
                total_score += 0.1
            else:
                # REF field exists but bookmark not found — might still be a valid cross-ref
                if clause52_bookmark_name and ref_target == clause52_bookmark_name:
                    print(f"PASS: Component 3 - REF field targets correct bookmark '{ref_target}' (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"PARTIAL: Component 3 - REF field exists but target '{ref_target}' not matched to Clause 5.2 heading (0.1 pts)")
                    total_score += 0.1
        elif has_fields:
            # Has field codes but couldn't parse REF — still indicates cross-reference attempt
            print(f"PARTIAL: Component 3 - Field codes found in 'see Clause 5.2' para but no parseable REF instruction (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 3 - No field codes in 'see Clause 5.2' paragraph")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Cross-reference field in 'as referenced in Section 3' paragraph (0.25 points)
    try:
        para_elem = doc.paragraphs[as_ref_section3_idx]._element
        ref_fields = find_ref_fields(para_elem)
        has_fields = has_field_codes(para_elem)

        if ref_fields:
            ref_target = ref_fields[0]
            target_para = all_bookmarks.get(ref_target)
            if target_para == section3_heading_idx:
                print(f"PASS: Component 4 - REF field '{ref_target}' in 'referenced in Section 3' para points to Section 3 heading (0.25 pts)")
                total_score += 0.25
            elif target_para is not None:
                print(f"PARTIAL: Component 4 - REF field '{ref_target}' points to para {target_para}, expected {section3_heading_idx} (0.1 pts)")
                total_score += 0.1
            else:
                if section3_bookmark_name and ref_target == section3_bookmark_name:
                    print(f"PASS: Component 4 - REF field targets correct bookmark '{ref_target}' (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"PARTIAL: Component 4 - REF field exists but target '{ref_target}' not matched to Section 3 heading (0.1 pts)")
                    total_score += 0.1
        elif has_fields:
            print(f"PARTIAL: Component 4 - Field codes found in 'referenced in Section 3' para but no parseable REF instruction (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 4 - No field codes in 'referenced in Section 3' paragraph")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
