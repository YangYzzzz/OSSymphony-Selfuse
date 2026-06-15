"""
Reward Script: Insert bibliography/references section at end of document
Task ID: writer_rd_044
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Bibliography entries exist after References heading
  Component 2 (0.30): Entry [1] - Smith et al. (2024) journal article
  Component 3 (0.25): Entry [2] - Johnson (2023) book
  Component 4 (0.30): Entry [3] - Lee and Wang (2025) conference paper
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_044'


def persist_app_state(domain):
    """Try to save any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in ("libreoffice_calc", "libreoffice_writer", "libreoffice_impress"):
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print("PERSIST: ctrl+s sent for " + domain)
        except Exception as e:
            print("PERSIST_WARN: save hook failed: " + str(e))


def check_entry_smith(bib_paragraphs):
    """Check for Smith et al. (2024) journal article entry. Returns score 0.0-0.30."""
    for bp in bib_paragraphs:
        bp_lower = bp.lower()
        if 'smith' in bp_lower and '2024' in bp_lower:
            pts = 0.15  # Base: author + year present
            if 'journal' in bp_lower or 'vol' in bp_lower or 'pp.' in bp_lower or 'no.' in bp_lower:
                pts += 0.10
                print("PASS: Component 2a — Smith 2024 has journal article markers")
            if 'martinez' in bp_lower or 'patel' in bp_lower:
                pts += 0.05
                print("PASS: Component 2b — Co-authors found")
            print("PASS: Component 2 — Smith et al. (2024) entry found (" + str(pts) + " pts)")
            return pts
    print("FAIL: Component 2 — No entry with 'Smith' and '2024' found")
    return 0.0


def check_entry_johnson(bib_paragraphs):
    """Check for Johnson (2023) book entry. Returns score 0.0-0.25."""
    for bp in bib_paragraphs:
        bp_lower = bp.lower()
        if 'johnson' in bp_lower and '2023' in bp_lower:
            pts = 0.15  # Base: author + year
            if 'press' in bp_lower or 'publisher' in bp_lower or 'cambridge' in bp_lower or 'storage' in bp_lower:
                pts += 0.10
                print("PASS: Component 3a — Johnson 2023 has book markers")
            print("PASS: Component 3 — Johnson (2023) entry found (" + str(pts) + " pts)")
            return pts
    print("FAIL: Component 3 — No entry with 'Johnson' and '2023' found")
    return 0.0


def check_entry_lee(bib_paragraphs):
    """Check for Lee and Wang (2025) conference paper entry. Returns score 0.0-0.30."""
    for bp in bib_paragraphs:
        bp_lower = bp.lower()
        if 'lee' in bp_lower and '2025' in bp_lower:
            pts = 0.15  # Base: author + year
            if 'wang' in bp_lower:
                pts += 0.05
                print("PASS: Component 4a — Wang co-author found")
            if 'proceedings' in bp_lower or 'conference' in bp_lower or 'singapore' in bp_lower:
                pts += 0.10
                print("PASS: Component 4b — Conference paper markers found")
            print("PASS: Component 4 — Lee and Wang (2025) entry found (" + str(pts) + " pts)")
            return pts
    print("FAIL: Component 4 — No entry with 'Lee' and '2025' found")
    return 0.0


def verify_task(file_path):
    """
    Verify bibliography section was added to the document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file " + file_path + ": " + str(e))
        print("REWARD: 0.0")
        return 0.0

    # Find the "References" heading and collect all paragraphs after it
    references_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'references' and para.style and 'Heading' in para.style.name:
            references_idx = i
            break

    if references_idx is None:
        print("FAIL: No 'References' heading found in document")
        print("REWARD: 0.0")
        return 0.0

    # Collect text of all non-empty paragraphs after the References heading
    bib_paragraphs = []
    for i in range(references_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text:
            bib_paragraphs.append(text)

    print("INFO: Found " + str(len(bib_paragraphs)) + " non-empty paragraphs after References heading")
    for idx, bp in enumerate(bib_paragraphs):
        print("  Entry " + str(idx) + ": " + bp[:100] + "...")

    # Component 1: Bibliography entries exist after References heading (0.15 points)
    # Initial doc has References heading with NO content below it
    try:
        if len(bib_paragraphs) >= 3:
            print("PASS: Component 1 — At least 3 bibliography entries found (" + str(len(bib_paragraphs)) + ") (0.15 pts)")
            total_score += 0.15
        elif len(bib_paragraphs) >= 1:
            print("PARTIAL: Component 1 — Only " + str(len(bib_paragraphs)) + " entries found, expected >= 3 (0.05 pts)")
            total_score += 0.05
        else:
            print("FAIL: Component 1 — No bibliography entries found after References heading")
    except Exception as e:
        print("ERROR: Component 1 — " + str(e))

    # Component 2: Entry [1] - Smith et al. (2024) journal article (0.30 points)
    try:
        if len(bib_paragraphs) > 0:
            total_score += check_entry_smith(bib_paragraphs)
    except Exception as e:
        print("ERROR: Component 2 — " + str(e))

    # Component 3: Entry [2] - Johnson (2023) book (0.25 points)
    try:
        if len(bib_paragraphs) > 0:
            total_score += check_entry_johnson(bib_paragraphs)
    except Exception as e:
        print("ERROR: Component 3 — " + str(e))

    # Component 4: Entry [3] - Lee and Wang (2025) conference paper (0.30 points)
    try:
        if len(bib_paragraphs) > 0:
            total_score += check_entry_lee(bib_paragraphs)
    except Exception as e:
        print("ERROR: Component 4 — " + str(e))

    final_score = round(min(total_score, 1.0), 2)
    print("\nScore: " + str(total_score) + "/1.0")
    print("REWARD: " + str(final_score))
    return final_score


# Persist app state before verification
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = WORKDIR + '/' + TASK_ID + '.docx'
if not os.path.exists(file_path):
    print("File not found: " + file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
