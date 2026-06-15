"""
Reward Script: Split cell D3 vertically into 3 rows for morning/afternoon/evening sessions
Task ID: writer_tm_037
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Table has 9 rows (2 added from original 7)
  Component 2 (0.35): Column D split region contains 3 distinct session labels
  Component 3 (0.25): Columns A,B,C,E are vertically merged across the 3 split rows
  Component 4 (0.20): Remaining data rows preserved correctly after split
"""

import os
from docx import Document
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_037'
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WNS}


def get_vmerge(tc):
    """Get vMerge status for a table cell element: 'restart', 'continue', or None."""
    tcPr = tc.find('w:tcPr', NS)
    if tcPr is None:
        return None
    vmerge = tcPr.find('w:vMerge', NS)
    if vmerge is None:
        return None
    val = vmerge.get(f'{{{WNS}}}val')
    if val is None:
        return 'continue'
    return val


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one table
    if len(doc.tables) < 1:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    # Component 1: Table row count increased to 9 (0.20 points)
    # Initial has 7 rows. Splitting D3 into 3 sub-rows adds 2 rows -> 9 total.
    try:
        if num_rows == 9:
            print(f"PASS: Component 1 — Table has 9 rows as expected (0.20 pts)")
            total_score += 0.20
        elif num_rows > 7:
            # Partial credit: rows were added but count is off
            print(f"PARTIAL: Component 1 — Table has {num_rows} rows (expected 9, initial was 7) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Table has {num_rows} rows (expected 9, initial was 7)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Column D split region contains session labels (0.35 points)
    # In the golden file, the 3 rows at the split region (rows 3,4,5) have
    # column 3 (D) containing: "Sessions", "Morning", "Afternoon"
    # GATE: requires 9+ rows (i.e., rows were actually added for the split)
    try:
        if num_rows >= 9:
            # Get the actual tc elements from XML for precise cell inspection
            split_d_texts = []
            split_d_no_vmerge = 0
            for ri in range(3, 6):
                row_elem = table.rows[ri]._tr
                tcs = row_elem.findall('w:tc', NS)
                if len(tcs) >= 4:
                    tc = tcs[3]  # column D (index 3)
                    vmerge_status = get_vmerge(tc)
                    # Column D cells should NOT be vertically merged (continue)
                    if vmerge_status != 'continue':
                        paras = tc.findall('w:p', NS)
                        text = ''
                        for p in paras:
                            for r in p.findall('.//w:t', NS):
                                if r.text:
                                    text += r.text
                        split_d_texts.append(text.strip())
                        split_d_no_vmerge += 1

            score_2 = 0.0
            # Check: we need 3 independent cells in column D within a 9-row table
            if split_d_no_vmerge >= 3:
                score_2 += 0.15
                print(f"PASS: Component 2a — Column D has {split_d_no_vmerge} independent cells in split region")
            else:
                print(f"FAIL: Component 2a — Column D has {split_d_no_vmerge} independent cells (expected 3)")

            # Check content: should have session-related labels (morning/afternoon/evening or Sessions)
            expected_time_labels = ['morning', 'afternoon', 'evening']
            time_labels_found = 0
            for text in split_d_texts:
                if any(label in text.lower() for label in expected_time_labels):
                    time_labels_found += 1

            if time_labels_found >= 2:
                score_2 += 0.20
                print(f"PASS: Component 2b — Found {time_labels_found} time-of-day labels: {split_d_texts} (0.20 pts)")
            elif time_labels_found >= 1:
                score_2 += 0.10
                print(f"PARTIAL: Component 2b — Found {time_labels_found} time-of-day labels: {split_d_texts} (0.10 pts)")
            else:
                print(f"FAIL: Component 2b — Found {time_labels_found} time-of-day labels: {split_d_texts}")

            if score_2 > 0:
                total_score += score_2
            print(f"  Component 2 total: {score_2:.2f}/0.35 pts")
        else:
            print(f"FAIL: Component 2 — Table has {num_rows} rows (need 9+ for split verification)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Columns A,B,C,E vertically merged across split rows (0.25 points)
    # In the golden file, columns 0,1,2,4 have vMerge=restart at row 3 and vMerge=continue at rows 4,5
    try:
        if num_rows >= 6:
            merge_cols = [0, 1, 2, 4]  # A, B, C, E
            merged_correctly = 0

            for col_idx in merge_cols:
                # Check all 3 merge conditions for this column
                checks_passed = 0
                # Row 3: should have vMerge=restart
                row3_tr = table.rows[3]._tr
                tcs3 = row3_tr.findall('w:tc', NS)
                if len(tcs3) > col_idx:
                    vm3 = get_vmerge(tcs3[col_idx])
                    if vm3 == 'restart':
                        checks_passed += 1

                # Rows 4,5: should have vMerge=continue
                for ri in range(4, 6):
                    row_tr = table.rows[ri]._tr
                    tcs = row_tr.findall('w:tc', NS)
                    if len(tcs) > col_idx:
                        vm = get_vmerge(tcs[col_idx])
                        if vm == 'continue':
                            checks_passed += 1

                if checks_passed == 3:
                    merged_correctly += 1

            if merged_correctly == 4:
                print(f"PASS: Component 3 — All 4 non-D columns are vertically merged (0.25 pts)")
                total_score += 0.25
            elif merged_correctly >= 2:
                partial = 0.25 * merged_correctly / 4
                print(f"PARTIAL: Component 3 — {merged_correctly}/4 columns correctly merged ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Only {merged_correctly}/4 columns correctly merged")
        else:
            print(f"FAIL: Component 3 — Not enough rows to check merges")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Remaining data rows preserved after split (0.20 points)
    # After the split region (rows 3-5), rows 6-8 should correspond to original rows 4-6
    # Original row 4: ['Tuesday', '11:00 - 12:30', 'Room 305', 'Data Engineering', 'Priya Patel']
    # Original row 5: ['Wednesday', '9:00 - 10:30', 'Hall A', 'AI Ethics Panel', 'Prof. Michael Torres']
    # Original row 6: ['Wednesday', '2:00 - 3:30', 'Main Stage', 'Closing Ceremony', 'Sarah Chen']
    try:
        if num_rows >= 9:
            expected_post_split = [
                ['Tuesday', '11:00 - 12:30', 'Room 305', 'Data Engineering', 'Priya Patel'],
                ['Wednesday', '9:00 - 10:30', 'Hall A', 'AI Ethics Panel', 'Prof. Michael Torres'],
                ['Wednesday', '2:00 - 3:30', 'Main Stage', 'Closing Ceremony', 'Sarah Chen'],
            ]
            rows_match = 0
            for idx, expected_row in enumerate(expected_post_split):
                actual_ri = 6 + idx
                if actual_ri < num_rows:
                    actual_texts = [table.rows[actual_ri].cells[ci].text.strip() for ci in range(min(5, num_cols))]
                    if actual_texts == expected_row:
                        rows_match += 1
                    else:
                        print(f"  Row {actual_ri} mismatch: expected {expected_row}, got {actual_texts}")

            if rows_match == 3:
                print(f"PASS: Component 4 — All 3 post-split data rows preserved correctly (0.20 pts)")
                total_score += 0.20
            elif rows_match >= 1:
                partial = 0.20 * rows_match / 3
                print(f"PARTIAL: Component 4 — {rows_match}/3 post-split rows correct ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — No post-split rows match expected data")
        else:
            print(f"FAIL: Component 4 — Not enough rows ({num_rows}) to check post-split data")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer (save unsaved GUI edits)
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
