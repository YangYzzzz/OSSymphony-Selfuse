"""
Initial Setup: Create Event Schedule document with a 5x7 table
Task ID: writer_tm_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Annual Technology Conference 2025", level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        "Below is the event schedule for the three-day Annual Technology Conference "
        "taking place at the Grand Convention Center, October 15-17, 2025."
    )
    intro.paragraph_format.space_after = Pt(12)

    # Create a 7-row x 5-column schedule table
    # Columns: Day | Time | Venue | Sessions | Speaker
    table = doc.add_table(rows=7, cols=5)
    table.style = "Table Grid"

    # Row 0: Headers
    headers = ["Day", "Time", "Venue", "Sessions", "Speaker"]
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Row 1
    data = [
        ["Monday", "9:00 - 10:30", "Hall A", "Opening Keynote", "Dr. Emily Zhang"],
        ["Monday", "11:00 - 12:30", "Room 201", "Cloud Architecture", "James Rivera"],
        ["Tuesday", "9:00 - 10:30", "Hall B", "Sessions", "Multiple"],
        ["Tuesday", "11:00 - 12:30", "Room 305", "Data Engineering", "Priya Patel"],
        ["Wednesday", "9:00 - 10:30", "Hall A", "AI Ethics Panel", "Prof. Michael Torres"],
        ["Wednesday", "2:00 - 3:30", "Main Stage", "Closing Ceremony", "Sarah Chen"],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx, col_idx).text = val

    # Add a note paragraph after the table
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "Note: All sessions include a 15-minute Q&A segment. "
        "Lunch break is from 12:30 to 1:30 PM daily in the Grand Dining Hall."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].italic = True

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
