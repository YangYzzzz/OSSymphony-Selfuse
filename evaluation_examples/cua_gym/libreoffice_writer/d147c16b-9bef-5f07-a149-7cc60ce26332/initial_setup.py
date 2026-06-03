"""
Initial Setup: Insert page breaks with landscape/portrait page styles
Task ID: writer_fs_027
Domain: libreoffice_writer

Creates a multi-page Writer document in portrait orientation containing:
- Title page, Introduction, Methodology, Results, Discussion, Conclusion sections
- 'Appendix A: Data Tables' section with wide data tables
- 'References' section at the end
All pages are portrait. No landscape sections or special page breaks.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_027'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_text(doc, text, space_after=Pt(6)):
    """Add a body paragraph with consistent formatting."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = space_after
    for run in para.runs:
        run.font.name = 'Liberation Serif'
        run.font.size = Pt(12)
    return para


def create_initial():
    doc = Document()

    # Set default page to portrait, standard US Letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # =====================
    # TITLE PAGE
    # =====================
    title = doc.add_heading('Annual Market Analysis Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Fiscal Year 2024-2025')
    run.font.size = Pt(16)
    run.font.name = 'Liberation Serif'
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Prepared by: Strategic Analytics Division\nGlobal Markets Research Group')
    run.font.size = Pt(12)
    run.font.name = 'Liberation Serif'

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_before = Pt(36)
    run = date_para.add_run('March 2025')
    run.font.size = Pt(14)
    run.font.name = 'Liberation Serif'

    # =====================
    # TABLE OF CONTENTS (placeholder)
    # =====================
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary ............................ 3',
        '2. Introduction ................................. 3',
        '3. Methodology .................................. 4',
        '4. Market Analysis Results ...................... 5',
        '5. Discussion ................................... 6',
        '6. Strategic Recommendations .................... 7',
        '7. Conclusion ................................... 7',
        'Appendix A: Data Tables ......................... 8',
        'References ....................................... 10',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.name = 'Liberation Serif'
            r.font.size = Pt(11)

    # =====================
    # EXECUTIVE SUMMARY (page 3)
    # =====================
    doc.add_page_break()
    doc.add_heading('1. Executive Summary', level=1)
    add_body_text(doc,
        'This report presents a comprehensive analysis of global market trends across '
        'key sectors during the fiscal year 2024-2025. Our research team examined data '
        'from 47 countries spanning six major economic regions to identify emerging '
        'patterns in consumer behavior, technology adoption, and financial market dynamics.')
    add_body_text(doc,
        'Key findings indicate a 12.3% increase in cross-border digital commerce, '
        'a significant shift toward sustainable investment portfolios, and accelerating '
        'adoption of artificial intelligence tools across mid-market enterprises. The '
        'technology sector maintained its dominant position with a combined market '
        'capitalization growth of 18.7% year-over-year.')
    add_body_text(doc,
        'Our strategic recommendations focus on three pillars: expanding digital '
        'infrastructure investments in Southeast Asian markets, diversifying supply '
        'chain partnerships to reduce single-region dependencies, and accelerating '
        'internal AI integration programs to maintain competitive advantages.')

    # =====================
    # INTRODUCTION (page 3-4)
    # =====================
    doc.add_heading('2. Introduction', level=1)
    add_body_text(doc,
        'The global economic landscape has undergone profound transformation during '
        'the period under review. Geopolitical tensions, evolving trade agreements, '
        'and rapid technological advancement have collectively reshaped market dynamics '
        'in ways that demand rigorous, data-driven analysis. This report aims to provide '
        'stakeholders with actionable insights derived from our proprietary research '
        'methodology and extensive data collection framework.')
    add_body_text(doc,
        'Our analysis encompasses multiple dimensions of market activity, including '
        'equity market performance, foreign exchange fluctuations, commodity pricing '
        'trends, and sector-specific growth indicators. Special attention has been '
        'given to emerging markets in the Asia-Pacific region, which continue to '
        'demonstrate outsized growth potential relative to established Western economies.')
    add_body_text(doc,
        'The research was conducted between January and December 2024, with supplementary '
        'data collected through Q1 2025 to capture recent developments. Our team of 23 '
        'analysts collaborated with regional offices in Singapore, Frankfurt, and '
        'São Paulo to ensure comprehensive geographic coverage and cultural context.')
    add_body_text(doc,
        'This report is structured to first outline our methodology, then present '
        'findings organized by sector and region, followed by an integrated discussion '
        'of cross-cutting themes and concluding with strategic recommendations for '
        'the upcoming fiscal year.')

    # =====================
    # METHODOLOGY (page 4-5)
    # =====================
    doc.add_page_break()
    doc.add_heading('3. Methodology', level=1)
    add_body_text(doc,
        'Our research methodology employs a multi-layered approach combining quantitative '
        'analysis of financial data with qualitative assessments from industry experts. '
        'The quantitative component utilizes proprietary algorithms to process daily '
        'market data from Bloomberg Terminal feeds, Reuters Eikon, and direct exchange '
        'API connections covering 12 major global exchanges.')

    doc.add_heading('3.1 Data Collection', level=2)
    add_body_text(doc,
        'Primary data sources include official exchange filings, central bank publications, '
        'and government statistical agencies from 47 countries. Secondary sources encompass '
        'industry reports from McKinsey Global Institute, World Economic Forum publications, '
        'and peer-reviewed academic journals in finance and economics.')
    add_body_text(doc,
        'Data validation was performed using a three-tier verification process: automated '
        'outlier detection algorithms flagged anomalous entries, cross-referencing between '
        'independent sources confirmed data integrity, and manual review by senior analysts '
        'resolved remaining discrepancies. This process resulted in a validated dataset '
        'comprising approximately 4.2 million data points.')

    doc.add_heading('3.2 Analytical Framework', level=2)
    add_body_text(doc,
        'The analytical framework is built on modern portfolio theory augmented with '
        'behavioral finance insights. We applied regression analysis, time-series '
        'decomposition, and machine learning classification models to identify patterns '
        'and predict near-term market trajectories. All statistical tests were conducted '
        'at the 95% confidence level unless otherwise specified.')
    add_body_text(doc,
        'Sector classification follows the Global Industry Classification Standard (GICS) '
        'with modifications to better capture emerging technology sub-sectors including '
        'quantum computing, space technology, and synthetic biology. Regional groupings '
        'align with IMF classifications but separate ASEAN nations into individual units '
        'given their growing individual significance.')

    # =====================
    # RESULTS (pages 5-6)
    # =====================
    doc.add_page_break()
    doc.add_heading('4. Market Analysis Results', level=1)

    doc.add_heading('4.1 Global Equity Markets', level=2)
    add_body_text(doc,
        'Global equity markets demonstrated resilience throughout 2024, with the MSCI '
        'World Index gaining 14.2% on a total-return basis. Developed markets outperformed '
        'emerging markets by 3.8 percentage points, largely driven by strong performance '
        'in the US technology sector and European healthcare companies.')
    add_body_text(doc,
        'Notable regional performances included Japan\'s Nikkei 225 reaching a 34-year '
        'high in February 2024, the Indian Sensex crossing 75,000 for the first time, '
        'and Brazil\'s Ibovespa recovering from mid-year volatility to finish the period '
        'up 8.9%. Chinese markets remained under pressure due to ongoing property sector '
        'concerns, with the CSI 300 declining 2.1% over the period.')

    doc.add_heading('4.2 Fixed Income and Currency Markets', level=2)
    add_body_text(doc,
        'Bond markets experienced significant volatility as central banks navigated the '
        'transition from tightening to easing monetary policy. The US 10-year Treasury '
        'yield fluctuated between 3.85% and 4.70%, reflecting shifting expectations for '
        'Federal Reserve rate cuts. European sovereign spreads narrowed as ECB policy '
        'provided stability, with the German-Italian 10-year spread compressing to 140 '
        'basis points by year end.')
    add_body_text(doc,
        'In currency markets, the US dollar maintained its strength against most major '
        'currencies, appreciating 4.3% on a trade-weighted basis. The Japanese yen '
        'reached a 34-year low against the dollar before intervention by the Bank of '
        'Japan stabilized the exchange rate. Emerging market currencies showed mixed '
        'performance, with the Indian rupee relatively stable while the Turkish lira '
        'continued its depreciation trend.')

    doc.add_heading('4.3 Technology Sector Deep Dive', level=2)
    add_body_text(doc,
        'The technology sector was the standout performer of the fiscal year. The '
        'NASDAQ Composite gained 28.6%, driven primarily by companies involved in '
        'artificial intelligence development and deployment. The combined market '
        'capitalization of the "Magnificent Seven" tech companies exceeded $13 trillion, '
        'representing approximately 30% of the S&P 500 total capitalization.')
    add_body_text(doc,
        'AI-related capital expenditure across major technology firms exceeded $200 '
        'billion in 2024, a 65% increase over the prior year. Semiconductor companies '
        'benefited from this demand surge, with NVIDIA\'s revenue growing 122% '
        'year-over-year. Cloud infrastructure spending accelerated as enterprises '
        'migrated workloads to support AI model training and inference capabilities.')

    # =====================
    # DISCUSSION (pages 6-7)
    # =====================
    doc.add_page_break()
    doc.add_heading('5. Discussion', level=1)
    add_body_text(doc,
        'The market dynamics observed during fiscal year 2024-2025 present both '
        'opportunities and challenges for investors and policymakers. The concentration '
        'of equity market returns in a small number of technology companies raises '
        'concerns about market breadth and potential vulnerability to sector-specific '
        'corrections. Historical analysis suggests that such concentrated markets have '
        'preceded periods of broader rotation.')
    add_body_text(doc,
        'The divergence between developed and emerging market performance highlights '
        'structural differences in monetary policy transmission and investor sentiment. '
        'While developed markets benefited from anticipated rate cuts and strong corporate '
        'earnings, emerging markets faced headwinds from capital outflows and domestic '
        'policy uncertainties. This gap is expected to narrow as global monetary conditions '
        'ease further in 2025.')
    add_body_text(doc,
        'Geopolitical risk remains a significant factor in portfolio construction. '
        'Supply chain diversification trends, often described as "friendshoring" or '
        '"nearshoring," are creating new investment opportunities in countries like '
        'Vietnam, Mexico, and India while potentially reducing efficiency in global '
        'manufacturing networks. The long-term implications for cost structures and '
        'profit margins across sectors warrant careful monitoring.')
    add_body_text(doc,
        'Environmental, Social, and Governance (ESG) considerations continued to '
        'influence capital allocation, albeit with growing debate about measurement '
        'methodology and actual impact. Sustainable investment funds attracted $42 '
        'billion in net inflows during the period, a 15% increase, though some '
        'strategies faced scrutiny over "greenwashing" concerns.')

    # =====================
    # STRATEGIC RECOMMENDATIONS (page 7)
    # =====================
    doc.add_heading('6. Strategic Recommendations', level=1)
    add_body_text(doc,
        'Based on our analysis, we recommend the following strategic actions for '
        'the upcoming fiscal year:')
    recommendations = [
        'Increase allocation to Southeast Asian equities by 3-5 percentage points, '
        'targeting Vietnam, Indonesia, and the Philippines as primary beneficiaries '
        'of supply chain diversification.',
        'Maintain overweight position in technology sector with emphasis on AI '
        'infrastructure companies while reducing exposure to consumer-facing tech '
        'firms trading above 40x forward earnings.',
        'Implement a barbell fixed-income strategy combining short-duration '
        'investment-grade bonds with select high-yield emerging market sovereign debt.',
        'Establish hedging positions against potential dollar weakness using options '
        'strategies as the Federal Reserve advances its easing cycle.',
        'Evaluate direct investment opportunities in renewable energy infrastructure '
        'across European and Asian markets where regulatory incentives are strongest.',
    ]
    for rec in recommendations:
        p = doc.add_paragraph(rec, style='List Number')
        for r in p.runs:
            r.font.name = 'Liberation Serif'
            r.font.size = Pt(12)

    # =====================
    # CONCLUSION (page 7-8)
    # =====================
    doc.add_heading('7. Conclusion', level=1)
    add_body_text(doc,
        'The fiscal year 2024-2025 has been characterized by technological transformation, '
        'monetary policy transitions, and evolving geopolitical alignments. Markets have '
        'demonstrated remarkable resilience in the face of these complex dynamics, though '
        'the concentration of returns and persistent valuations in certain sectors warrant '
        'vigilance.')
    add_body_text(doc,
        'Looking ahead, we anticipate a more balanced market environment as rate cuts '
        'support broader equity participation and emerging markets benefit from improved '
        'capital flows. The AI investment cycle remains in its early stages and will '
        'continue to reshape sector dynamics and competitive landscapes across industries. '
        'Our team will continue to monitor these developments and provide updated guidance '
        'through quarterly supplements to this annual report.')

    # =====================
    # APPENDIX A: DATA TABLES (page 8-9)
    # =====================
    doc.add_page_break()
    doc.add_heading('Appendix A: Data Tables', level=1)
    add_body_text(doc,
        'The following tables present detailed data supporting the analysis in this report.')

    # Table 1: Regional Market Performance
    add_body_text(doc, 'Table A.1: Regional Equity Market Performance (FY 2024-2025)',
                  space_after=Pt(4))
    t1_headers = ['Region', 'Index', 'Start Value', 'End Value', 'Change (%)',
                  'Volatility', 'P/E Ratio', 'Dividend Yield (%)']
    t1_data = [
        ['North America', 'S&P 500', '4,769.83', '5,450.21', '+14.3%', '13.2', '21.8', '1.42'],
        ['North America', 'NASDAQ Comp.', '15,011.35', '19,305.60', '+28.6%', '17.8', '34.2', '0.68'],
        ['Europe', 'STOXX 600', '476.31', '528.49', '+10.9%', '11.5', '14.6', '3.21'],
        ['Europe', 'FTSE 100', '7,733.24', '8,194.17', '+6.0%', '10.8', '12.1', '3.85'],
        ['Asia Pacific', 'Nikkei 225', '33,464.17', '39,756.43', '+18.8%', '15.4', '18.9', '1.78'],
        ['Asia Pacific', 'Hang Seng', '17,047.39', '16,592.28', '-2.7%', '18.9', '8.7', '4.12'],
        ['Emerging', 'MSCI EM', '987.42', '1,038.91', '+5.2%', '14.7', '12.3', '2.94'],
        ['Latin America', 'Ibovespa', '134,185', '146,117', '+8.9%', '16.3', '8.4', '5.67'],
    ]

    table1 = doc.add_table(rows=1 + len(t1_data), cols=len(t1_headers))
    table1.style = 'Table Grid'
    for i, h in enumerate(t1_headers):
        cell = table1.cell(0, i)
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = 'Liberation Sans'
    for row_idx, row_data in enumerate(t1_data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table1.cell(row_idx, col_idx)
            cell.text = val
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9)
                r.font.name = 'Liberation Sans'

    add_body_text(doc, '')  # spacer

    # Table 2: Sector Performance Breakdown
    add_body_text(doc, 'Table A.2: Sector Performance Breakdown (GICS Classification)',
                  space_after=Pt(4))
    t2_headers = ['Sector', 'Market Cap ($B)', 'YoY Change (%)', 'Revenue Growth (%)',
                  'Earnings Growth (%)', 'Avg P/E', 'Fwd P/E', 'Weight in S&P 500 (%)']
    t2_data = [
        ['Information Technology', '14,832', '+24.7', '+12.8', '+18.3', '32.4', '28.1', '31.2'],
        ['Healthcare', '5,245', '+8.4', '+6.2', '+9.1', '19.7', '17.3', '12.8'],
        ['Financials', '4,987', '+12.1', '+8.5', '+14.6', '13.2', '11.8', '13.1'],
        ['Consumer Discretionary', '4,321', '+15.3', '+7.9', '+11.2', '24.8', '21.5', '10.4'],
        ['Communication Services', '3,876', '+19.8', '+9.4', '+22.7', '20.3', '17.9', '8.9'],
        ['Industrials', '3,654', '+9.7', '+5.8', '+7.3', '18.6', '16.4', '8.7'],
        ['Consumer Staples', '2,987', '+4.2', '+3.1', '+5.8', '21.3', '19.7', '6.2'],
        ['Energy', '2,543', '-3.8', '+1.2', '-8.4', '11.4', '10.9', '4.1'],
        ['Utilities', '1,876', '+6.5', '+4.7', '+8.9', '16.8', '15.2', '2.5'],
        ['Materials', '1,654', '+2.3', '+1.8', '+3.4', '14.9', '13.6', '2.1'],
        ['Real Estate', '1,432', '+7.8', '+5.3', '+6.7', '35.2', '29.8', '2.4'],
    ]

    table2 = doc.add_table(rows=1 + len(t2_data), cols=len(t2_headers))
    table2.style = 'Table Grid'
    for i, h in enumerate(t2_headers):
        cell = table2.cell(0, i)
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = 'Liberation Sans'
    for row_idx, row_data in enumerate(t2_data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table2.cell(row_idx, col_idx)
            cell.text = val
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9)
                r.font.name = 'Liberation Sans'

    add_body_text(doc, '')  # spacer

    # Table 3: Currency Exchange Rates
    add_body_text(doc, 'Table A.3: Major Currency Exchange Rates vs USD',
                  space_after=Pt(4))
    t3_headers = ['Currency', 'Code', 'Start Rate', 'End Rate', 'Change (%)',
                  'High', 'Low', '30-Day Vol (%)']
    t3_data = [
        ['Euro', 'EUR/USD', '1.1050', '1.0830', '-2.0', '1.1275', '1.0620', '7.8'],
        ['British Pound', 'GBP/USD', '1.2730', '1.2615', '-0.9', '1.3045', '1.2305', '8.2'],
        ['Japanese Yen', 'USD/JPY', '141.04', '151.32', '+7.3', '161.95', '140.25', '10.4'],
        ['Swiss Franc', 'USD/CHF', '0.8415', '0.8680', '+3.2', '0.9225', '0.8335', '7.1'],
        ['Canadian Dollar', 'USD/CAD', '1.3240', '1.3585', '+2.6', '1.3845', '1.3155', '5.9'],
        ['Australian Dollar', 'AUD/USD', '0.6825', '0.6540', '-4.2', '0.6940', '0.6350', '9.3'],
        ['Chinese Yuan', 'USD/CNY', '7.0980', '7.2470', '+2.1', '7.3125', '7.0425', '3.8'],
        ['Indian Rupee', 'USD/INR', '83.15', '83.52', '+0.4', '84.08', '82.90', '2.6'],
    ]

    table3 = doc.add_table(rows=1 + len(t3_data), cols=len(t3_headers))
    table3.style = 'Table Grid'
    for i, h in enumerate(t3_headers):
        cell = table3.cell(0, i)
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = 'Liberation Sans'
    for row_idx, row_data in enumerate(t3_data, 1):
        for col_idx, val in enumerate(row_data):
            cell = table3.cell(row_idx, col_idx)
            cell.text = val
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9)
                r.font.name = 'Liberation Sans'

    # =====================
    # REFERENCES (page 10)
    # =====================
    doc.add_page_break()
    doc.add_heading('References', level=1)
    references = [
        'Acemoglu, D., & Restrepo, P. (2024). "Artificial Intelligence, Automation, and Work." '
        'American Economic Review, 114(3), 488-524.',
        'Bank for International Settlements. (2024). Annual Economic Report 2024. Basel: BIS Press.',
        'Bloomberg Intelligence. (2024). Global Equity Market Outlook: Q4 2024 Update. '
        'New York: Bloomberg LP.',
        'Damodaran, A. (2024). "Equity Risk Premiums: Determinants, Estimation and Implications." '
        'Stern School of Business Working Paper.',
        'European Central Bank. (2024). Financial Stability Review, November 2024. '
        'Frankfurt: ECB Publications.',
        'International Monetary Fund. (2024). World Economic Outlook: October 2024. '
        'Washington, DC: IMF.',
        'McKinsey Global Institute. (2024). "The State of AI in 2024: Generative AI\'s Breakout Year." '
        'McKinsey & Company.',
        'Morgan Stanley Research. (2024). "2025 Global Strategy Outlook: The Great Rotation." '
        'New York: Morgan Stanley.',
        'OECD. (2024). Economic Outlook, Volume 2024 Issue 2. Paris: OECD Publishing.',
        'World Bank Group. (2024). Global Economic Prospects, June 2024. Washington, DC: World Bank.',
        'World Economic Forum. (2024). Global Risks Report 2024, 19th Edition. Geneva: WEF.',
        'Yellen, J. L. (2024). "Remarks on the State of the US Economy and Global Markets." '
        'US Department of the Treasury, Press Release, September 15, 2024.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for r in p.runs:
            r.font.name = 'Liberation Serif'
            r.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
