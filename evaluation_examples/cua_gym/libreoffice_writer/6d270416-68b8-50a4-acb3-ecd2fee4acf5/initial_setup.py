"""
Initial Setup: Financial report with currency amounts in dollar format
Task ID: writer_frd_023
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_023'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Quarterly Financial Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies Inc. - Q4 2025')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph('')  # spacer

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    # Amounts 1, 2, 3
    doc.add_paragraph(
        'Meridian Technologies concluded Q4 2025 with total revenue of $1,234,567.89, '
        'representing a significant increase over the previous quarter. Operating expenses '
        'were held at $845,230.50, resulting in a net operating income that exceeded '
        'projections by $52,400.00.'
    )

    # --- Revenue Breakdown ---
    doc.add_heading('Revenue Breakdown', level=1)
    # Amounts 4, 5, 6, 7
    doc.add_paragraph(
        'Product sales contributed $678,450.25 to the quarterly total, while '
        'subscription services generated $312,890.75. Consulting engagements brought in '
        '$156,780.00, and licensing fees accounted for $86,446.89.'
    )

    # --- Operating Expenses ---
    doc.add_heading('Operating Expenses', level=1)
    # Amounts 8, 9, 10
    doc.add_paragraph(
        'Salaries and benefits represented the largest expense category at $425,600.00. '
        'Marketing and advertising expenditures totaled $187,350.25, while research and '
        'development costs reached $132,280.25.'
    )

    # --- Capital Expenditures ---
    doc.add_heading('Capital Expenditures', level=1)
    # Amounts 11, 12
    doc.add_paragraph(
        'The company invested $275,000.00 in new server infrastructure during Q4. '
        'Additionally, office renovation costs amounted to $48,500.00, which were '
        'allocated across the facilities budget.'
    )

    # --- Regional Performance ---
    doc.add_heading('Regional Performance', level=1)
    # Amounts 13, 14
    doc.add_paragraph(
        'The North American division generated $534,200.50 in revenue, maintaining its '
        'position as the top-performing region. The European division contributed '
        '$389,120.75 despite unfavorable currency exchange conditions during the quarter.'
    )

    # --- Outlook ---
    doc.add_heading('Outlook and Projections', level=1)
    # Amounts 15, 16
    doc.add_paragraph(
        'Management projects Q1 2026 revenue to reach $1,450,000.00 based on the current '
        'sales pipeline. A contingency reserve of $75,000.00 has been set aside to '
        'address potential supply chain disruptions in the Asia-Pacific region.'
    )

    doc.add_paragraph('')  # spacer

    # --- Signature ---
    sig = doc.add_paragraph()
    sig.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = sig.add_run('Prepared by: Elena Vasquez, CFO')
    run.font.italic = True
    run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
