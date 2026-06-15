"""
Reward Script: Add a visible top border line to the footer area
Task ID: writer_fs_094
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Footer paragraph has a top border element
  Component 2 (0.25): Border style is "single" (solid line) with color #404040
  Component 3 (0.25): Border size is ~1 pt (w:sz=8, i.e. 8 half-points)
  Component 4 (0.15): Border spacing is ~0.3 cm above footer content (w:space ~ 7-11 pt)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_094'

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section with a footer
    if len(doc.sections) == 0:
        print("FAIL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    # We check all sections; if ANY section's footer has the border, we consider it.
    # The task says "on all pages", so ideally all sections should have it.
    # But for simplicity, check the first section's footer (most common case).
    section = doc.sections[0]
    footer = section.footer

    if not footer.paragraphs:
        print("FAIL: Footer has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # Find the footer paragraph (the one with page number content)
    target_para = None
    for para in footer.paragraphs:
        target_para = para
        break  # Use first paragraph

    if target_para is None:
        print("FAIL: No target paragraph found in footer")
        print("REWARD: 0.0")
        return 0.0

    # Extract border info from the paragraph's pPr/pBdr/top element
    pPr = target_para._element.find('w:pPr', ns)
    pBdr = None
    top_border = None

    if pPr is not None:
        pBdr = pPr.find('w:pBdr', ns)
    if pBdr is not None:
        top_border = pBdr.find('w:top', ns)

    # Component 1: Footer paragraph has a top border element (0.35 points)
    try:
        if top_border is not None:
            val_attr = top_border.get(qn('w:val'))
            # Ensure it's not "none" or "nil" which would mean no visible border
            if val_attr and val_attr not in ('none', 'nil'):
                print(f"PASS: Component 1 -- Top border element exists with val='{val_attr}' (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 1 -- Top border val is '{val_attr}' (no visible border)")
        else:
            print("FAIL: Component 1 -- No top border element found in footer paragraph")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Border style is "single" and color is #404040 (0.25 points)
    try:
        if top_border is not None:
            val_attr = top_border.get(qn('w:val'))
            color_attr = top_border.get(qn('w:color'))

            style_ok = val_attr == 'single'
            # Allow color match with some tolerance: exact 404040 or close variants
            color_ok = False
            if color_attr:
                color_upper = color_attr.upper().lstrip('#')
                color_ok = color_upper == '404040'

            if style_ok and color_ok:
                print(f"PASS: Component 2 -- Border style='single', color='#{color_attr}' (0.25 pts)")
                total_score += 0.25
            elif style_ok:
                # Partial: style correct but color wrong -- give half
                print(f"PARTIAL: Component 2 -- Style='single' correct, but color='{color_attr}' (expected 404040) (0.125 pts)")
                total_score += 0.125
            elif color_ok:
                print(f"PARTIAL: Component 2 -- Color correct, but style='{val_attr}' (expected single) (0.125 pts)")
                total_score += 0.125
            else:
                print(f"FAIL: Component 2 -- style='{val_attr}' (expected single), color='{color_attr}' (expected 404040)")
        else:
            print("FAIL: Component 2 -- No top border element to check style/color")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Border size is ~1 pt (w:sz=8 means 8 half-points = 1pt) (0.25 points)
    try:
        if top_border is not None:
            sz_attr = top_border.get(qn('w:sz'))
            if sz_attr:
                sz_val = int(sz_attr)
                # 1 pt = 8 half-points. Allow range 6-10 for tolerance.
                if sz_val == 8:
                    print(f"PASS: Component 3 -- Border size={sz_val} half-pts (1.0 pt exactly) (0.25 pts)")
                    total_score += 0.25
                elif 6 <= sz_val <= 10:
                    print(f"PARTIAL: Component 3 -- Border size={sz_val} half-pts (~{sz_val/8:.2f} pt, close to 1 pt) (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 3 -- Border size={sz_val} half-pts ({sz_val/8:.2f} pt), expected ~8 (1 pt)")
            else:
                print("FAIL: Component 3 -- No sz attribute on top border")
        else:
            print("FAIL: Component 3 -- No top border element to check size")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Border spacing ~0.3 cm above footer (w:space in points, 0.3 cm ~ 8.5 pt) (0.15 points)
    try:
        if top_border is not None:
            space_attr = top_border.get(qn('w:space'))
            if space_attr:
                space_val = int(space_attr)
                # 0.3 cm = ~8.5 pt. Allow range 5-14 for reasonable tolerance.
                if 5 <= space_val <= 14:
                    print(f"PASS: Component 4 -- Border spacing={space_val} pt (~{space_val * 0.0353:.2f} cm) (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 4 -- Border spacing={space_val} pt, expected ~7-11 (0.3 cm)")
            else:
                # No space attribute; some implementations may default to 0
                print("FAIL: Component 4 -- No space attribute on top border")
        else:
            print("FAIL: Component 4 -- No top border element to check spacing")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
