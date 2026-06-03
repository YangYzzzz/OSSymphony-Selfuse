"""
Initial Setup: Writer document with footer containing page numbers, no border.
Task ID: writer_fs_094
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_094'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_to_footer(section):
    """Add page number field code to footer paragraph."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_text = fp.add_run("Page ")
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # PAGE field code
    r1 = fp.add_run()
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    instr.set(qn('xml:space'), 'preserve')
    r2._element.append(instr)

    r3 = fp.add_run()
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading("Quarterly Performance Review - Q1 2025", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Prepared by the Human Resources Department")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    doc.add_paragraph()  # spacer

    # Section 1
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "This report presents the quarterly performance review for all departments "
        "during Q1 2025. Overall, the company has met 87% of its key performance "
        "indicators, representing a 4% improvement over Q4 2024. Revenue growth "
        "exceeded projections by 2.3%, driven primarily by strong performance in the "
        "Enterprise Solutions and Cloud Infrastructure divisions."
    )
    doc.add_paragraph(
        "Employee satisfaction scores remained stable at 4.2 out of 5.0, while "
        "voluntary attrition decreased from 8.1% to 6.7%. The Learning & Development "
        "team launched 14 new training modules, achieving an average completion rate "
        "of 78% across all departments."
    )

    # Section 2
    doc.add_heading("Department Highlights", level=2)

    doc.add_heading("Engineering", level=3)
    doc.add_paragraph(
        "The Engineering team delivered 23 sprint cycles with a velocity increase of "
        "15%. Key achievements include the successful migration of the payment "
        "processing microservice to Kubernetes, reducing deployment times by 40%. "
        "Technical debt reduction initiatives addressed 156 legacy issues, bringing "
        "the backlog below the quarterly target of 200 items."
    )

    doc.add_heading("Sales & Marketing", level=3)
    doc.add_paragraph(
        "New client acquisition reached 47 accounts, exceeding the target of 40. "
        "The average deal size grew from $32,500 to $38,200, attributed to the "
        "revised pricing strategy implemented in January. Digital marketing campaigns "
        "generated 12,400 qualified leads through a combination of content marketing, "
        "webinar series, and targeted LinkedIn outreach."
    )

    doc.add_heading("Customer Success", level=3)
    doc.add_paragraph(
        "Net Promoter Score improved from 62 to 71, placing the company in the top "
        "quartile of industry benchmarks. The support team resolved 94% of tickets "
        "within the defined SLA windows. Three new self-service knowledge base "
        "articles were published weekly, reducing Tier 1 ticket volume by 18%."
    )

    # Section 3 - enough content for page 2
    doc.add_heading("Financial Overview", level=2)
    doc.add_paragraph(
        "Total revenue for Q1 2025 reached $14.7 million, up from $13.2 million in "
        "Q4 2024. Operating expenses were held at $11.1 million, resulting in an "
        "operating margin of 24.5%. Capital expenditures totaled $1.8 million, "
        "primarily allocated to data center expansion and new development tooling."
    )
    doc.add_paragraph(
        "The accounts receivable aging report shows 91% of invoices collected within "
        "45 days, an improvement over the prior quarter's 86%. Cash reserves stand at "
        "$8.3 million, providing a comfortable runway for planned Q2 investments in "
        "product development and market expansion initiatives."
    )

    doc.add_heading("Strategic Initiatives", level=2)
    doc.add_paragraph(
        "The AI Integration Program entered Phase 2, with pilot deployments across "
        "three enterprise clients. Early feedback indicates a 30% reduction in manual "
        "data processing time. The International Expansion Committee finalized "
        "partnership agreements with distributors in Germany and Japan, with targeted "
        "market entry scheduled for Q3 2025."
    )
    doc.add_paragraph(
        "Sustainability efforts continue with the Green Office Initiative achieving "
        "a 22% reduction in energy consumption year-over-year. The company received "
        "ISO 14001 environmental management certification in February, strengthening "
        "our competitive positioning in environmentally conscious markets."
    )

    doc.add_heading("Recommendations", level=2)
    doc.add_paragraph(
        "Based on the Q1 results, the leadership team recommends increasing "
        "investment in the Enterprise Solutions division by 12%, accelerating the "
        "hiring plan for the Cloud Infrastructure team from 8 to 12 positions, and "
        "expanding the customer success onboarding program to include all accounts "
        "with annual contract values exceeding $50,000."
    )

    # Footer with page numbers (NO border - that's the task)
    add_page_number_to_footer(section)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
