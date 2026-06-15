"""
Reward Script: Incident Report Form in LibreOffice Writer
Task ID: writer_wf_034
Domain: libreoffice_writer
Scoring:
  Component 1: Title 'INCIDENT REPORT FORM' centered and bold (0.15)
  Component 2: Reporter Information table 4x2 (Name, Department, Date, Time) (0.20)
  Component 3: Incident Details table 3x2 (Location, Type placeholder, Description) (0.20)
  Component 4: Witnesses table 4x3 (header + 3 rows) (0.20)
  Component 5: Numbered list for Immediate Actions Taken (0.10)
  Component 6: Supervisor Review with signature and date lines (0.15)
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_034'


def persist_app_state(domain):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph texts and table info for analysis
    paragraphs = doc.paragraphs
    tables = doc.tables

    # Component 1: Title 'INCIDENT REPORT FORM' centered and bold (0.15 points)
    try:
        title_found = False
        for p in paragraphs:
            text = p.text.strip().upper()
            if 'INCIDENT REPORT FORM' in text:
                # Check centered alignment
                is_centered = (p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                # Check bold - at least one run containing the title text must be bold
                has_bold = any(r.bold for r in p.runs if r.text.strip())
                if is_centered and has_bold:
                    print(f"PASS: Component 1 — Title 'INCIDENT REPORT FORM' is centered and bold (0.15 pts)")
                    total_score += 0.15
                    title_found = True
                else:
                    print(f"FAIL: Component 1 — Title found but centered={is_centered}, bold={has_bold}")
                    title_found = True
                break
        if not title_found:
            print("FAIL: Component 1 — Title 'INCIDENT REPORT FORM' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Reporter Information table (0.20 points)
    # Must be a table with at least 4 rows and 2 cols, containing Name, Department, Date, Time labels
    try:
        reporter_table_found = False
        required_labels = {'name', 'department', 'date', 'time'}
        for table in tables:
            if len(table.rows) >= 4 and len(table.columns) >= 2:
                # Get all first-column cell texts
                col0_texts = set()
                for row in table.rows:
                    cell_text = row.cells[0].text.strip().lower().rstrip(':')
                    col0_texts.add(cell_text)
                if required_labels.issubset(col0_texts):
                    print(f"PASS: Component 2 — Reporter Information table found with {len(table.rows)} rows, labels: {col0_texts} (0.20 pts)")
                    total_score += 0.20
                    reporter_table_found = True
                    break
        if not reporter_table_found:
            print("FAIL: Component 2 — Reporter Information table with Name/Department/Date/Time not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Incident Details table (0.20 points)
    # Must have Location, Type (with dropdown placeholder text), and Description rows
    try:
        incident_table_found = False
        for table in tables:
            col0_texts_lower = []
            for row in table.rows:
                col0_texts_lower.append(row.cells[0].text.strip().lower().rstrip(':'))
            has_location = any('location' in t for t in col0_texts_lower)
            has_type = any('type' in t for t in col0_texts_lower)
            has_description = any('description' in t for t in col0_texts_lower)
            if has_location and has_type and has_description:
                # Check for type dropdown placeholder text in the type row
                type_placeholder_found = False
                for row in table.rows:
                    if 'type' in row.cells[0].text.strip().lower():
                        cell_text = row.cells[1].text.strip().lower() if len(row.cells) > 1 else ''
                        if 'select' in cell_text or 'injury' in cell_text or 'near miss' in cell_text or 'dropdown' in cell_text or '[' in cell_text:
                            type_placeholder_found = True
                        break
                if type_placeholder_found:
                    print(f"PASS: Component 3 — Incident Details table with Location/Type(placeholder)/Description (0.20 pts)")
                    total_score += 0.20
                else:
                    # Partial: table structure correct but no placeholder
                    print(f"PARTIAL: Component 3 — Incident Details table found but Type row lacks dropdown placeholder text (0.10 pts)")
                    total_score += 0.10
                incident_table_found = True
                break
        if not incident_table_found:
            print("FAIL: Component 3 — Incident Details table with Location/Type/Description not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Witnesses table with 3 data rows (0.20 points)
    # Header row with Name, Contact, Statement + 3 empty rows = 4 rows, 3 cols
    try:
        witness_table_found = False
        for table in tables:
            if len(table.columns) >= 3:
                header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
                has_name = any('name' in t for t in header_texts)
                has_contact = any('contact' in t for t in header_texts)
                has_statement = any('statement' in t for t in header_texts)
                if has_name and has_contact and has_statement:
                    # Count data rows (rows after header)
                    data_rows = len(table.rows) - 1
                    if data_rows >= 3:
                        print(f"PASS: Component 4 — Witnesses table with header + {data_rows} data rows (0.20 pts)")
                        total_score += 0.20
                    else:
                        print(f"PARTIAL: Component 4 — Witnesses table found but only {data_rows} data rows (need 3) (0.10 pts)")
                        total_score += 0.10
                    witness_table_found = True
                    break
        if not witness_table_found:
            print("FAIL: Component 4 — Witnesses table with Name/Contact/Statement header not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Numbered list for Immediate Actions Taken (0.10 points)
    # Look for section heading + numbered items (1. 2. 3. etc.)
    try:
        actions_section_found = False
        numbered_items = 0
        in_actions = False
        for p in paragraphs:
            text = p.text.strip()
            if 'immediate actions taken' in text.lower():
                in_actions = True
                actions_section_found = True
                continue
            if in_actions:
                # Check if this looks like a numbered item
                import re
                if re.match(r'^\d+[\.\)]\s*', text):
                    numbered_items += 1
                # Stop if we hit another section header (bold non-empty text that isn't numbered)
                elif text and not re.match(r'^\d+[\.\)]\s*', text) and any(r.bold for r in p.runs if r.text.strip()):
                    break

        if actions_section_found and numbered_items >= 3:
            print(f"PASS: Component 5 — Immediate Actions Taken with {numbered_items} numbered items (0.10 pts)")
            total_score += 0.10
        elif actions_section_found:
            print(f"FAIL: Component 5 — Section found but only {numbered_items} numbered items (need >=3)")
        else:
            print("FAIL: Component 5 — 'Immediate Actions Taken' section not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Supervisor Review section with signature and date lines (0.15 points)
    try:
        has_supervisor_section = False
        has_signature_line = False
        has_date_line = False
        in_supervisor = False
        for p in paragraphs:
            text = p.text.strip()
            lower_text = text.lower()
            # Mark section start only for a heading-like paragraph (short, bold, "supervisor review")
            if not in_supervisor and 'supervisor' in lower_text and 'review' in lower_text:
                has_supervisor_section = True
                in_supervisor = True
                continue
            if in_supervisor:
                if 'signature' in lower_text and ('___' in text or '____' in text):
                    has_signature_line = True
                if 'date' in lower_text and ('___' in text or '____' in text):
                    has_date_line = True

        if has_supervisor_section and has_signature_line and has_date_line:
            print(f"PASS: Component 6 — Supervisor Review with signature and date lines (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 6 — supervisor={has_supervisor_section}, signature={has_signature_line}, date={has_date_line}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
