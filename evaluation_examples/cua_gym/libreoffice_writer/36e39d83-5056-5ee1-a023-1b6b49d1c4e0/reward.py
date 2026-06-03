"""
Reward Script: Promote level 3 and level 4 list items to level 2 in a multi-level outline
Task ID: writer_list_053
Domain: libreoffice_writer
Scoring:
  Component 1: Former level-3 items (ilvl=2) are now at level 2 (ilvl=1)  — 0.4 pts
                Items: "Competitor benchmarking", "Customer segmentation", "Organic growth initiatives"
  Component 2: Former level-4 items (ilvl=3) are now at level 2 (ilvl=1)  — 0.3 pts
                Items: "Pricing comparison", "Content marketing plan"
  Component 3: Full two-level structure: "Strategic Planning" at ilvl=0,
                all 7 other items at ilvl=1                                 — 0.3 pts
Total: 1.0

Initial state (ilvl values): 0, 1, 2, 3, 2, 1, 2, 3
Golden state (ilvl values):  0, 1, 1, 1, 1, 1, 1, 1
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_list_053'


def get_ilvl(para):
    """Extract the list indentation level (ilvl) from paragraph XML. Returns None if not in a list."""
    numPr = para._element.find(qn('w:pPr') + '/' + qn('w:numPr'))
    if numPr is None:
        return None
    ilvl_elem = numPr.find(qn('w:ilvl'))
    if ilvl_elem is None:
        return None
    val = ilvl_elem.get(qn('w:val'))
    if val is None:
        return None
    return int(val)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires promoting all level-3 (ilvl=2) and level-4 (ilvl=3) items
    to level-2 (ilvl=1) in the numbered outline list in deep_outline.docx.
    """
    total_score = 0.0

    # Precondition gate: load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify all expected paragraph texts are present
    expected_texts = [
        "Strategic Planning",
        "Market Analysis",
        "Competitor benchmarking",
        "Pricing comparison",
        "Customer segmentation",
        "Growth Strategy",
        "Organic growth initiatives",
        "Content marketing plan",
    ]
    actual_texts = [p.text for p in doc.paragraphs]
    if not all(t in actual_texts for t in expected_texts):
        print(f"CRITICAL: Document is missing expected paragraphs. Found: {actual_texts}")
        print("REWARD: 0.0")
        return 0.0

    # Build a mapping of text -> ilvl for convenience
    text_to_ilvl = {}
    for para in doc.paragraphs:
        if para.text in expected_texts:
            text_to_ilvl[para.text] = get_ilvl(para)

    print(f"Paragraph ilvl mapping: {text_to_ilvl}")

    # Component 1: Former level-3 items (originally ilvl=2) are now at ilvl=1 (0.4 points)
    # These are: "Competitor benchmarking", "Customer segmentation", "Organic growth initiatives"
    former_level3_items = [
        "Competitor benchmarking",
        "Customer segmentation",
        "Organic growth initiatives",
    ]
    try:
        level3_promoted = []
        level3_failed = []
        for item in former_level3_items:
            actual_ilvl = text_to_ilvl.get(item)
            if actual_ilvl == 1:
                level3_promoted.append(item)
            else:
                level3_failed.append(f"{item!r} has ilvl={actual_ilvl} (expected 1)")

        if len(level3_promoted) == len(former_level3_items):
            print(f"PASS: Component 1 — All 3 former level-3 items promoted to ilvl=1: {level3_promoted} (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Not all former level-3 items at ilvl=1. Failures: {level3_failed}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Former level-4 items (originally ilvl=3) are now at ilvl=1 (0.3 points)
    # These are: "Pricing comparison", "Content marketing plan"
    former_level4_items = [
        "Pricing comparison",
        "Content marketing plan",
    ]
    try:
        level4_promoted = []
        level4_failed = []
        for item in former_level4_items:
            actual_ilvl = text_to_ilvl.get(item)
            if actual_ilvl == 1:
                level4_promoted.append(item)
            else:
                level4_failed.append(f"{item!r} has ilvl={actual_ilvl} (expected 1)")

        if len(level4_promoted) == len(former_level4_items):
            print(f"PASS: Component 2 — Both former level-4 items promoted to ilvl=1: {level4_promoted} (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Not all former level-4 items at ilvl=1. Failures: {level4_failed}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Full two-level structure verification (0.3 points)
    # "Strategic Planning" must be at ilvl=0, and ALL other 7 items must be at ilvl=1
    # This confirms the document has exactly 2 levels, not just partial promotion
    try:
        root_item = "Strategic Planning"
        root_ilvl = text_to_ilvl.get(root_item)

        non_root_items = [t for t in expected_texts if t != root_item]
        all_level2 = all(text_to_ilvl.get(t) == 1 for t in non_root_items)
        root_correct = (root_ilvl == 0)

        if root_correct and all_level2:
            print(f"PASS: Component 3 — Full two-level structure confirmed: '{root_item}' at ilvl=0, all 7 other items at ilvl=1 (0.3 pts)")
            total_score += 0.3
        else:
            issues = []
            if not root_correct:
                issues.append(f"'{root_item}' has ilvl={root_ilvl} (expected 0)")
            for t in non_root_items:
                lvl = text_to_ilvl.get(t)
                if lvl != 1:
                    issues.append(f"'{t}' has ilvl={lvl} (expected 1)")
            print(f"FAIL: Component 3 — Two-level structure not achieved. Issues: {issues}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Run verification against the task file
file_path = f'{WORKDIR}/Desktop/deep_outline.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
