"""
Initial Setup: Create a four-level numbered outline document
Task ID: writer_list_053
Domain: libreoffice_writer

Creates /home/user/Desktop/deep_outline.docx with a 4-level numbered list:
  Level 1: Strategic Planning
  Level 2: Market Analysis, Growth Strategy
  Level 3: Competitor benchmarking, Customer segmentation, Organic growth initiatives
  Level 4: Pricing comparison, Content marketing plan
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'deep_outline'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_num_pr(num_id: int, ilvl: int) -> OxmlElement:
    """Create a <w:numPr> element for list numbering."""
    num_pr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    num_id_el = OxmlElement('w:numId')
    num_id_el.set(qn('w:val'), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    return num_pr


def add_abstract_num(doc: Document) -> int:
    """
    Add an abstractNum definition for a 4-level decimal outline (1. / 1.1 / 1.1.1 / 1.1.1.1).
    Returns the abstractNumId value.
    """
    numbering_part = doc.part.numbering_part
    numbering = numbering_part._element

    # Choose a fresh abstractNumId
    existing_ids = [int(el.get(qn('w:abstractNumId')))
                    for el in numbering.findall(qn('w:abstractNum'))]
    abstract_num_id = max(existing_ids, default=-1) + 1

    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))

    # multiLevelType
    multi_level_type = OxmlElement('w:multiLevelType')
    multi_level_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level_type)

    # Define 4 levels (ilvl 0-3)
    for ilvl in range(4):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'decimal')
        lvl.append(num_fmt)

        # Build lvlText: "1." / "%1.%2." / "%1.%2.%3." / "%1.%2.%3.%4."
        lvl_text = OxmlElement('w:lvlText')
        text_val = '.'.join(f'%{i+1}' for i in range(ilvl + 1)) + '.'
        lvl_text.set(qn('w:val'), text_val)
        lvl.append(lvl_text)

        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind_left = 360 * (ilvl + 1)
        ind.set(qn('w:left'), str(ind_left))
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl.append(pPr)

        abstract_num.append(lvl)

    # Insert before any existing <w:num> elements
    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        numbering.insert(list(numbering).index(first_num), abstract_num)
    else:
        numbering.append(abstract_num)

    return abstract_num_id


def add_num_instance(doc: Document, abstract_num_id: int) -> int:
    """
    Create a <w:num> instance referencing the abstract definition.
    Returns the numId value.
    """
    numbering_part = doc.part.numbering_part
    numbering = numbering_part._element

    existing_num_ids = [int(el.get(qn('w:numId')))
                        for el in numbering.findall(qn('w:num'))]
    num_id = max(existing_num_ids, default=0) + 1

    num_el = OxmlElement('w:num')
    num_el.set(qn('w:numId'), str(num_id))

    abstract_num_id_el = OxmlElement('w:abstractNumId')
    abstract_num_id_el.set(qn('w:val'), str(abstract_num_id))
    num_el.append(abstract_num_id_el)

    numbering.append(num_el)
    return num_id


def add_list_item(doc: Document, text: str, level: int, num_id: int):
    """Add a paragraph as a numbered list item at the given 0-based level."""
    para = doc.add_paragraph()
    run = para.add_run(text)

    pPr = para._p.get_or_add_pPr()
    pPr.append(make_num_pr(num_id, level))
    return para


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Ensure numbering part exists (add a dummy list to trigger creation, then remove)
    dummy = doc.add_paragraph('x', style='List Number')
    dummy._element.getparent().remove(dummy._element)

    # Add abstract num definition and get a concrete numId
    abstract_num_id = add_abstract_num(doc)
    num_id = add_num_instance(doc, abstract_num_id)

    # Build the 4-level outline structure
    # Level 1 (ilvl=0): "Strategic Planning"
    add_list_item(doc, 'Strategic Planning', 0, num_id)
    # Level 2 (ilvl=1): "Market Analysis"
    add_list_item(doc, 'Market Analysis', 1, num_id)
    # Level 3 (ilvl=2): "Competitor benchmarking"
    add_list_item(doc, 'Competitor benchmarking', 2, num_id)
    # Level 4 (ilvl=3): "Pricing comparison"
    add_list_item(doc, 'Pricing comparison', 3, num_id)
    # Level 3 (ilvl=2): "Customer segmentation"
    add_list_item(doc, 'Customer segmentation', 2, num_id)
    # Level 2 (ilvl=1): "Growth Strategy"
    add_list_item(doc, 'Growth Strategy', 1, num_id)
    # Level 3 (ilvl=2): "Organic growth initiatives"
    add_list_item(doc, 'Organic growth initiatives', 2, num_id)
    # Level 4 (ilvl=3): "Content marketing plan"
    add_list_item(doc, 'Content marketing plan', 3, num_id)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
