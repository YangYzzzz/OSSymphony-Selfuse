"""
Initial Setup: Business report document with Author = 'Sarah Mitchell'
Task ID: writer_biz_074
Domain: libreoffice_writer

Creates a realistic business report document. No macros or form buttons exist.
Document properties Author set to 'Sarah Mitchell'.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set document core properties — Author = 'Sarah Mitchell'
    core = doc.core_properties
    core.author = 'Sarah Mitchell'
    core.title = 'Q1 2025 Regional Sales Performance Report'
    core.subject = 'Quarterly Sales Analysis'
    core.keywords = 'sales, quarterly, performance, 2025'

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = hp.add_run('Pinnacle Solutions Inc. — Confidential')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.italic = True

    # --- Footer ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_f = fp.add_run('Page ')
    run_f.font.size = Pt(9)
    r1 = fp.add_run()
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    # === TITLE PAGE ===
    title_para = doc.add_heading('Q1 2025 Regional Sales Performance Report', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_s = subtitle.add_run('Pinnacle Solutions Inc.')
    run_s.font.size = Pt(16)
    run_s.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_d = date_para.add_run('Prepared: March 31, 2025')
    run_d.font.size = Pt(12)
    run_d.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_a = author_para.add_run('Author: Sarah Mitchell, VP of Sales Operations')
    run_a.font.size = Pt(11)

    doc.add_page_break()

    # === EXECUTIVE SUMMARY ===
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents the Q1 2025 sales performance analysis for Pinnacle Solutions Inc. '
        'across all four operating regions. Overall revenue increased by 12.3% compared to Q4 2024, '
        'driven primarily by strong growth in the Western and Southern regions. The Eastern region '
        'maintained steady performance, while the Northern region showed signs of recovery after '
        'the downturn experienced in the previous two quarters.'
    )
    doc.add_paragraph(
        'Key highlights include the successful launch of the Enterprise Cloud Suite, which contributed '
        '$2.4M in new recurring revenue, and the expansion of our strategic partnership with GlobalTech '
        'Systems, resulting in three major enterprise deals totaling $1.8M.'
    )

    # === REGIONAL BREAKDOWN ===
    doc.add_heading('2. Regional Performance Breakdown', level=1)

    # Table: Regional Summary
    doc.add_heading('2.1 Revenue by Region', level=2)
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    headers = ['Region', 'Q1 2025 Revenue', 'Q4 2024 Revenue', 'Change ($)', 'Change (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['Western', '$4,215,300', '$3,780,100', '+$435,200', '+11.5%'],
        ['Eastern', '$3,890,750', '$3,820,400', '+$70,350', '+1.8%'],
        ['Southern', '$3,156,200', '$2,645,800', '+$510,400', '+19.3%'],
        ['Northern', '$2,478,600', '$2,310,500', '+$168,100', '+7.3%'],
        ['Total', '$13,740,850', '$12,556,800', '+$1,184,050', '+9.4%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val
    # Bold the total row
    for c in range(5):
        for run in table.cell(5, c).paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph('')  # spacer

    # === TOP PERFORMERS ===
    doc.add_heading('2.2 Top Performing Sales Representatives', level=2)
    doc.add_paragraph(
        'The following representatives exceeded their quarterly targets by more than 20%:'
    )

    top_performers = [
        'Jessica Ramirez (Western) — $892,400 closed, 142% of target',
        'David Park (Southern) — $765,200 closed, 138% of target',
        'Amara Okafor (Eastern) — $654,800 closed, 131% of target',
        'Ryan Kowalski (Northern) — $512,300 closed, 127% of target',
        'Priya Sharma (Western) — $498,700 closed, 124% of target',
    ]
    for perf in top_performers:
        doc.add_paragraph(perf, style='List Bullet')

    # === PRODUCT ANALYSIS ===
    doc.add_heading('3. Product Line Analysis', level=1)
    doc.add_paragraph(
        'The Enterprise Cloud Suite launch in January 2025 has been the standout performer this '
        'quarter. Initial adoption rates exceeded projections by 34%, with 47 new enterprise '
        'customers onboarded during the quarter. The average deal size for Enterprise Cloud Suite '
        'was $51,200, compared to $32,800 for our legacy Professional Suite.'
    )

    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    headers2 = ['Product Line', 'Q1 Revenue', 'Units Sold', 'Avg Deal Size']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    prod_data = [
        ['Enterprise Cloud Suite', '$2,406,400', '47', '$51,200'],
        ['Professional Suite', '$4,822,800', '147', '$32,800'],
        ['Starter Package', '$3,945,650', '438', '$9,010'],
        ['Custom Integrations', '$2,566,000', '28', '$91,643'],
    ]
    for r, row_data in enumerate(prod_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_paragraph('')

    # === STRATEGIC INITIATIVES ===
    doc.add_heading('4. Strategic Initiatives & Outlook', level=1)
    doc.add_paragraph(
        'Looking ahead to Q2 2025, we anticipate continued momentum driven by:'
    )
    initiatives = [
        'Expansion of the Enterprise Cloud Suite into the Asia-Pacific market, with pilot programs '
        'scheduled in Singapore and Tokyo offices.',
        'Launch of the AI-powered Analytics Dashboard add-on, targeting existing Professional Suite '
        'customers for upsell opportunities.',
        'Hiring of 15 additional sales development representatives to support the growing pipeline '
        'in the Southern region.',
        'Renewal of the GlobalTech Systems partnership agreement with expanded scope covering '
        'professional services and implementation support.',
    ]
    for init in initiatives:
        doc.add_paragraph(init, style='List Number')

    # === RISKS & CHALLENGES ===
    doc.add_heading('5. Risks and Challenges', level=1)
    doc.add_paragraph(
        'Several factors could impact Q2 performance and should be monitored closely:'
    )
    doc.add_paragraph(
        'Competitive pressure from NovaSoft\'s recently launched CloudFirst platform, which offers '
        'similar functionality at a 15% lower price point. Our product differentiation strategy '
        'must be clearly communicated to the sales team.'
    )
    doc.add_paragraph(
        'Supply chain delays in hardware provisioning for on-premise deployments may affect '
        'delivery timelines for three pending enterprise contracts worth a combined $1.2M.'
    )
    doc.add_paragraph(
        'The Northern region continues to face challenges with staff retention. Two senior account '
        'executives departed in February, and their pipeline requires redistribution.'
    )

    # === CONCLUSION ===
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        'Q1 2025 represents a strong start to the fiscal year for Pinnacle Solutions Inc. The '
        'successful launch of the Enterprise Cloud Suite and the robust performance of the Western '
        'and Southern regions position us well for continued growth. Management recommends '
        'accelerating investment in the Asia-Pacific expansion and the AI Analytics Dashboard to '
        'capitalize on current market momentum.'
    )

    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_before = Pt(36)
    run_sig = sig_para.add_run('Sarah Mitchell')
    run_sig.bold = True
    run_sig.font.size = Pt(11)
    sig_title = doc.add_paragraph('VP of Sales Operations')
    sig_title.paragraph_format.space_before = Pt(0)
    sig_date = doc.add_paragraph('March 31, 2025')
    sig_date.paragraph_format.space_before = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
