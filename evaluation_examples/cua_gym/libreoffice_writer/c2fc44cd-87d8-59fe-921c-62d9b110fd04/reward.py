"""
Reward Script: Set proportional line spacing at 120% and 0.5 cm first-line indent
               for all body text paragraphs in the newsletter article.
Task ID: writer_para_068
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All 4 body text paragraphs have proportional line spacing at 120%
                          (line_spacing=1.2, line_spacing_rule=MULTIPLE)
  Component 2 (0.3 pts): All 4 body text paragraphs have first_line_indent ~0.5 cm
  Component 3 (0.2 pts): All 4 body text paragraphs have BOTH correct line spacing AND
                          correct first-line indent simultaneously (full task completion bonus),
                          AND heading paragraphs remain unmodified
Total: 1.0
"""

import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_LINE_SPACING

WORKDIR = '/home/user'
TASK_ID = 'writer_para_068'
FILE_PATH = f'{WORKDIR}/Desktop/alumni_newsletter.docx'

# Index positions of body text paragraphs (must be modified)
BODY_PARA_INDICES = [2, 3, 5, 6]
# Index positions of heading paragraphs (must remain unchanged)
HEADING_PARA_INDICES = [0, 1, 4]

# Expected values
EXPECTED_LINE_SPACING = 1.2  # 120% proportional
EXPECTED_FIRST_LINE_INDENT_EMU = 179705  # ~0.5 cm in EMU (actual value set by LibreOffice)
INDENT_TOLERANCE_EMU = 10000  # ±tolerance for EMU comparison (~0.028 cm)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have expected paragraph count
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected at least 7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 4 body text paragraphs have proportional line spacing at 120% (0.5 points)
    # Fails on initial_env (ls=None for all body paragraphs), passes on golden_env (ls=1.2, MULTIPLE)
    try:
        body_ls_pass_count = 0
        for idx in BODY_PARA_INDICES:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            ls = pf.line_spacing
            ls_rule = pf.line_spacing_rule
            # line_spacing=1.2 with rule=MULTIPLE (5) indicates 120% proportional spacing
            if (ls is not None and abs(float(ls) - EXPECTED_LINE_SPACING) < 0.05
                    and ls_rule == WD_LINE_SPACING.MULTIPLE):
                body_ls_pass_count += 1
                print(f"  PASS-C1: Para {idx} ({para.text[:30]!r}) ls={ls}, rule={ls_rule}")
            else:
                print(f"  FAIL-C1: Para {idx} ({para.text[:30]!r}) ls={ls}, rule={ls_rule}"
                      f" (expected ls=1.2, rule=MULTIPLE)")

        if body_ls_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 1 — All {len(BODY_PARA_INDICES)} body paragraphs have "
                  f"proportional 120% line spacing (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Only {body_ls_pass_count}/{len(BODY_PARA_INDICES)} "
                  f"body paragraphs have correct line spacing")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 4 body text paragraphs have first_line_indent ~0.5 cm (0.3 points)
    # Fails on initial_env (fli=None for all body paragraphs), passes on golden_env (~179705 EMU)
    try:
        body_fli_pass_count = 0
        for idx in BODY_PARA_INDICES:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            fli = pf.first_line_indent
            # Check EMU value is within tolerance of 0.5 cm (~180000 EMU)
            if (fli is not None and
                    abs(int(fli) - EXPECTED_FIRST_LINE_INDENT_EMU) <= INDENT_TOLERANCE_EMU):
                body_fli_pass_count += 1
                fli_cm = int(fli) / 360000
                print(f"  PASS-C2: Para {idx} ({para.text[:30]!r}) fli={fli} ({fli_cm:.3f} cm)")
            else:
                fli_cm = (int(fli) / 360000) if fli is not None else None
                print(f"  FAIL-C2: Para {idx} ({para.text[:30]!r}) fli={fli} "
                      f"({fli_cm} cm) (expected ~{EXPECTED_FIRST_LINE_INDENT_EMU} EMU / 0.5 cm)")

        if body_fli_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 2 — All {len(BODY_PARA_INDICES)} body paragraphs have "
                  f"first-line indent ~0.5 cm (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Only {body_fli_pass_count}/{len(BODY_PARA_INDICES)} "
                  f"body paragraphs have correct first-line indent")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Full completion bonus — all body paragraphs have BOTH correct line spacing
    # AND correct indent, AND heading paragraphs remain unchanged (0.2 points)
    # This compound check fails on initial_env (body paragraphs lack both properties) and
    # passes only on golden_env (all requirements simultaneously satisfied).
    try:
        compound_pass_count = 0
        for idx in BODY_PARA_INDICES:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            ls = pf.line_spacing
            ls_rule = pf.line_spacing_rule
            fli = pf.first_line_indent
            has_ls = (ls is not None and abs(float(ls) - EXPECTED_LINE_SPACING) < 0.05
                      and ls_rule == WD_LINE_SPACING.MULTIPLE)
            has_fli = (fli is not None and
                       abs(int(fli) - EXPECTED_FIRST_LINE_INDENT_EMU) <= INDENT_TOLERANCE_EMU)
            if has_ls and has_fli:
                compound_pass_count += 1

        heading_violations = []
        for idx in HEADING_PARA_INDICES:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            ls = pf.line_spacing
            fli = pf.first_line_indent
            if ls is not None or fli is not None:
                heading_violations.append(idx)
                print(f"  FAIL-C3: Heading para {idx} ({para.text[:30]!r}) was unexpectedly "
                      f"modified: ls={ls}, fli={fli}")
        headings_unchanged = (len(heading_violations) == 0)

        all_body_correct = (compound_pass_count == len(BODY_PARA_INDICES))
        if all_body_correct and headings_unchanged:
            print(f"PASS: Component 3 — All body paragraphs fully formatted AND headings "
                  f"unchanged (0.2 pts)")
            total_score += 0.2
        else:
            if not all_body_correct:
                print(f"FAIL: Component 3 — Only {compound_pass_count}/{len(BODY_PARA_INDICES)} "
                      f"body paragraphs pass the compound check (ls AND fli)")
            if not headings_unchanged:
                print(f"FAIL: Component 3 — One or more heading paragraphs were modified")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
