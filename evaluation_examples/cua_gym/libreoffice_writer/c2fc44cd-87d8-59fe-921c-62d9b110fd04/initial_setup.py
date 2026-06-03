"""
Initial Setup: Set proportional line spacing at 120% and 0.5 cm first-line indent for all body text paragraphs
Task ID: writer_para_068
Domain: libreoffice_writer

Creates alumni_newsletter.docx with 7 paragraphs (2 Heading 2s, 1 Heading 1, 4 body text)
in its pre-task state: no proportional line spacing, no first-line indent on body text.
"""

import os
import shlex
import shutil
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_068'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP_PATH = f'{WORKDIR}/Desktop/alumni_newsletter.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Remove default empty paragraph that Document() creates
    # (we will add paragraphs manually)
    # Note: Document() creates one empty paragraph; we'll clear it
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Paragraph 1: 'University Alumni Newsletter — Spring 2025' (Heading 1)
    h1 = doc.add_heading('University Alumni Newsletter \u2014 Spring 2025', level=1)

    # Paragraph 2: 'Message from the Dean' (Heading 2)
    h2a = doc.add_heading('Message from the Dean', level=2)

    # Paragraph 3: Body text (NO proportional line spacing, NO first-line indent)
    p3 = doc.add_paragraph(
        'Dear Alumni, it is my great pleasure to share with you the remarkable achievements '
        'of our university community over the past semester. Your continued support has been '
        'instrumental in advancing our mission of academic excellence.'
    )
    # Ensure no special line spacing or indent is applied
    pf3 = p3.paragraph_format
    pf3.line_spacing = None
    pf3.first_line_indent = None

    # Paragraph 4: Body text (NO proportional line spacing, NO first-line indent)
    p4 = doc.add_paragraph(
        'This spring, we welcomed the largest incoming class in university history, with 2,800 '
        'new students representing 65 countries. Our scholarship fund, to which many of you have '
        'generously contributed, provided $4.2 million in financial aid.'
    )
    pf4 = p4.paragraph_format
    pf4.line_spacing = None
    pf4.first_line_indent = None

    # Paragraph 5: 'Campus Development Update' (Heading 2)
    h2b = doc.add_heading('Campus Development Update', level=2)

    # Paragraph 6: Body text (NO proportional line spacing, NO first-line indent)
    p6 = doc.add_paragraph(
        'Construction of the new Science and Innovation Center is on track for completion in '
        'September 2025. The 50,000 square foot facility will house state-of-the-art laboratories '
        'for biotechnology, materials science, and quantum computing research.'
    )
    pf6 = p6.paragraph_format
    pf6.line_spacing = None
    pf6.first_line_indent = None

    # Paragraph 7: Body text (NO proportional line spacing, NO first-line indent)
    p7 = doc.add_paragraph(
        'The renovated Student Commons, funded by the Class of 1995 gift, opened in January to '
        'overwhelmingly positive student feedback.'
    )
    pf7 = p7.paragraph_format
    pf7.line_spacing = None
    pf7.first_line_indent = None

    # Save to canonical task_id path
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also copy to Desktop location as referenced in the task instruction
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)
    shutil.copy(OUTPUT, DESKTOP_PATH)
    print(f'Copied to Desktop: {DESKTOP_PATH}')

    # GUI-ready startup: open the Desktop file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DESKTOP_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
