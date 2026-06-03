"""
Initial Setup: New Employee Training Manual with 15 onboarding steps as plain paragraphs
Task ID: writer_hr_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document title ---
    title = doc.add_heading('New Employee Training Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / intro
    intro = doc.add_paragraph()
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = intro.add_run('Acme Corp - Human Resources Department')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph(
        'Welcome to Acme Corp! This training manual outlines the essential onboarding steps '
        'that every new employee must complete during their first 90 days. Please follow each '
        'step carefully and check in with your manager or HR representative if you have questions.'
    )

    doc.add_paragraph('')  # blank separator

    section_heading = doc.add_heading('Onboarding Procedures', level=1)

    doc.add_paragraph(
        'The following procedures must be completed in order. Each step includes a brief '
        'description of what is expected.'
    )

    # --- 15 onboarding steps as plain paragraphs (Normal style) ---
    # Each step: a step-title paragraph followed by a description paragraph
    # ALL in Normal style - no numbering, no indentation

    steps = [
        (
            'Complete HR Orientation Paperwork',
            'Visit the HR office on the 3rd floor to sign your employment contract, tax forms '
            '(W-4 and state withholding), direct deposit authorization, and emergency contact '
            'information. Bring two forms of government-issued identification.'
        ),
        (
            'Attend Company Welcome Session',
            'Join the scheduled welcome session held every Monday at 9:00 AM in Conference Room B. '
            'The session covers company history, organizational structure, core values, and the '
            'employee code of conduct. Light refreshments will be provided.'
        ),
        (
            'Set Up Workstation and Equipment',
            'Coordinate with the IT Help Desk (ext. 4500) to receive your laptop, monitors, '
            'keyboard, and headset. Ensure your docking station is configured and all peripherals '
            'are functioning before proceeding to the next step.'
        ),
        (
            'Configure Corporate Email and Calendar',
            'Log into your Acme Corp email account using the temporary credentials provided by IT. '
            'Set up your email signature following the company template, configure calendar sharing '
            'with your team, and subscribe to relevant department distribution lists.'
        ),
        (
            'Enroll in Benefits Program',
            'Access the benefits portal at benefits.acmecorp.com within your first 30 days to '
            'select your health insurance plan, dental and vision coverage, life insurance options, '
            'and 401(k) contribution percentage. Late enrollment may result in a waiting period.'
        ),
        (
            'Complete Mandatory Safety Training',
            'Register for the online safety training module through the Learning Management System. '
            'This covers fire evacuation routes, first aid station locations, hazardous material '
            'handling procedures, and workplace ergonomics. A passing score of 80% is required.'
        ),
        (
            'Review Information Security Policies',
            'Read and acknowledge the company information security handbook available on the '
            'intranet. Topics include password management, data classification levels, acceptable '
            'use of company devices, phishing awareness, and incident reporting procedures.'
        ),
        (
            'Meet Your Assigned Buddy',
            'Your department manager will introduce you to your onboarding buddy during your first '
            'week. Schedule a 30-minute coffee chat to learn about team norms, recommended lunch '
            'spots, parking tips, and any unwritten office customs that will help you settle in.'
        ),
        (
            'Tour the Office Facilities',
            'Your buddy or office coordinator will take you on a guided tour covering the main '
            'work areas, break rooms, fitness center, mail room, supply closet, and reserved '
            'parking areas. Note the locations of emergency exits and AED devices on each floor.'
        ),
        (
            'Obtain Building Access Badge',
            'Visit the Security Office in the lobby with your employee ID confirmation email. '
            'A photo will be taken for your access badge, which grants entry to the building, '
            'your assigned floor, and designated restricted areas based on your role.'
        ),
        (
            'Set Up Project Management Tools',
            'Request access to Jira, Confluence, and Slack from your team lead. Join your '
            "project's Slack channels, review the team wiki on Confluence for current sprint "
            'goals, and familiarize yourself with the ticketing workflow used in your department.'
        ),
        (
            'Schedule One-on-One with Direct Manager',
            'Book a 45-minute introductory meeting with your direct manager within the first two '
            'weeks. Prepare to discuss your role expectations, initial project assignments, '
            'performance review timelines, and any professional development goals you may have.'
        ),
        (
            'Complete Compliance and Ethics Training',
            'Finish the annual compliance training module covering anti-harassment policies, '
            'diversity and inclusion guidelines, conflict of interest disclosure, and whistleblower '
            'protection protocols. Certification must be completed within 60 days of hire.'
        ),
        (
            'Register for Professional Development',
            'Browse the training catalog on the Learning Management System and enroll in at '
            'least one skill-building course relevant to your role. Popular options include '
            'leadership fundamentals, advanced Excel, presentation skills, and time management.'
        ),
        (
            'Submit 30-Day Check-In Self-Assessment',
            'At the end of your first month, complete the self-assessment form available on the '
            'HR portal. Reflect on your onboarding experience, note any challenges encountered, '
            'and outline your goals for the next 60 days. Your manager will review this with you.'
        ),
    ]

    for step_title, step_desc in steps:
        # Step title as a plain paragraph (Normal style, no numbering)
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(step_title)
        run_title.bold = True

        # Step description as a plain paragraph (Normal style, no indentation)
        doc.add_paragraph(step_desc)

    # Closing note
    doc.add_paragraph('')
    closing = doc.add_paragraph()
    run_close = closing.add_run(
        'Please ensure all steps are completed within your first 90 days. Contact the HR '
        'department at hr@acmecorp.com or extension 2100 for any questions.'
    )
    run_close.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
