"""
Reward Script: Fitness class schedule for Peak Performance Gym
Task ID: writer_wf_078
Domain: libreoffice_writer
Scoring:
  C1 (0.15) - Title: 20pt bold centered with gym name
  C2 (0.20) - Table structure: >= 15 rows, 7 cols, time slots as rows, days as cols
  C3 (0.20) - At least 15 classes distributed in the table
  C4 (0.20) - Bold morning classes / italic evening classes
  C5 (0.10) - Instructor names section below table
  C6 (0.15) - Footer with address and phone
"""

import os
import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_078'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title — 20pt bold centered with gym name (0.15 pts)
    try:
        title_found = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Check if this paragraph mentions the gym name and looks like a title
            if 'peak performance' in text.lower() and ('gym' in text.lower() or 'fitness' in text.lower()):
                # Check centered
                is_centered = para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                # Check bold and 20pt in runs
                has_bold = False
                has_20pt = False
                for run in para.runs:
                    if run.font.bold:
                        has_bold = True
                    if run.font.size and abs(run.font.size.pt - 20.0) < 1.0:
                        has_20pt = True
                if is_centered and has_bold and has_20pt:
                    print(f"PASS: Component 1 — Title found: '{text[:60]}', centered, bold, 20pt (0.15 pts)")
                    total_score += 0.15
                    title_found = True
                else:
                    print(f"PARTIAL: Component 1 — Title text found but formatting off: centered={is_centered}, bold={has_bold}, 20pt={has_20pt}")
                    # Partial credit if text is right but formatting is not all there
                    if has_bold or has_20pt or is_centered:
                        total_score += 0.05
                title_found = True
                break
        if not title_found:
            print("FAIL: Component 1 — No title paragraph with 'Peak Performance Gym' found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table structure — at least 15 rows, 7 columns, time+day headers (0.20 pts)
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)

            # Check header row has day names
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            days_expected = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday']
            days_found = sum(1 for d in days_expected if any(d in h for h in header_cells))

            # Check time column has time slots
            time_pattern = re.compile(r'\d{1,2}:\d{2}\s*(AM|PM|am|pm)', re.IGNORECASE)
            time_count = 0
            for row in table.rows[1:]:
                cell_text = row.cells[0].text.strip()
                if time_pattern.search(cell_text):
                    time_count += 1

            structure_ok = num_rows >= 15 and num_cols >= 7 and days_found >= 5 and time_count >= 10
            if structure_ok:
                print(f"PASS: Component 2 — Table: {num_rows} rows x {num_cols} cols, {days_found} days, {time_count} time slots (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — Table: {num_rows} rows x {num_cols} cols, {days_found}/6 days, {time_count} time slots (need >=15 rows, 7 cols, 5 days, 10 times)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: At least 15 classes distributed in the table (0.20 pts)
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No tables found")
        else:
            table = doc.tables[0]
            class_count = 0
            for ri, row in enumerate(table.rows):
                if ri == 0:
                    continue  # skip header
                for ci, cell in enumerate(row.cells):
                    if ci == 0:
                        continue  # skip time column
                    text = cell.text.strip()
                    if text and len(text) > 1:
                        class_count += 1

            if class_count >= 15:
                print(f"PASS: Component 3 — {class_count} classes found in table (>= 15) (0.20 pts)")
                total_score += 0.20
            elif class_count >= 10:
                partial = 0.10
                print(f"PARTIAL: Component 3 — {class_count} classes found (need >= 15, giving {partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Only {class_count} classes found in table (need >= 15)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Bold morning classes / italic evening classes (0.20 pts)
    # Morning: rows with AM times (before 12:00 PM) should have bold runs
    # Evening: rows with PM times >= 5:00 PM should have italic runs
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 4 — No tables found")
        else:
            table = doc.tables[0]
            morning_bold_count = 0
            morning_total = 0
            evening_italic_count = 0
            evening_total = 0

            for ri, row in enumerate(table.rows):
                if ri == 0:
                    continue
                time_text = row.cells[0].text.strip().upper()
                # Determine if morning or evening
                is_morning = 'AM' in time_text
                # Parse hour for PM evening detection
                is_evening = False
                pm_match = re.search(r'(\d{1,2}):\d{2}\s*PM', time_text, re.IGNORECASE)
                if pm_match:
                    hour = int(pm_match.group(1))
                    if hour >= 5 or hour == 12:
                        # 5PM+ or 12PM is afternoon, but task says "evening"
                        # Based on golden: 5PM+ are italic
                        if hour >= 5 and hour != 12:
                            is_evening = True

                for ci, cell in enumerate(row.cells):
                    if ci == 0:
                        continue
                    text = cell.text.strip()
                    if not text:
                        continue

                    cell_bold = False
                    cell_italic = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                if run.font.bold:
                                    cell_bold = True
                                if run.font.italic:
                                    cell_italic = True

                    if is_morning:
                        morning_total += 1
                        if cell_bold:
                            morning_bold_count += 1
                    elif is_evening:
                        evening_total += 1
                        if cell_italic:
                            evening_italic_count += 1

            # Need at least some morning classes bold and some evening classes italic
            morning_ok = morning_total > 0 and morning_bold_count >= max(1, morning_total * 0.5)
            evening_ok = evening_total > 0 and evening_italic_count >= max(1, evening_total * 0.5)

            if morning_ok and evening_ok:
                print(f"PASS: Component 4 — Morning bold: {morning_bold_count}/{morning_total}, Evening italic: {evening_italic_count}/{evening_total} (0.20 pts)")
                total_score += 0.20
            elif morning_ok or evening_ok:
                print(f"PARTIAL: Component 4 — Morning bold: {morning_bold_count}/{morning_total}, Evening italic: {evening_italic_count}/{evening_total} (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 — Morning bold: {morning_bold_count}/{morning_total}, Evening italic: {evening_italic_count}/{evening_total}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Instructor names section below table (0.10 pts)
    try:
        instructor_paras = []
        # Look for paragraphs after the table that mention instructor-like content
        found_instructor_header = False
        for para in doc.paragraphs:
            text = para.text.strip().lower()
            if 'instructor' in text:
                found_instructor_header = True
                continue
            if found_instructor_header and text:
                # Instructor entries typically have name patterns (colon or dash separator)
                if ':' in para.text or '-' in para.text or len(text.split()) >= 2:
                    instructor_paras.append(para.text.strip())

        if found_instructor_header and len(instructor_paras) >= 3:
            print(f"PASS: Component 5 — Instructor section found with {len(instructor_paras)} entries (0.10 pts)")
            total_score += 0.10
        elif found_instructor_header:
            print(f"PARTIAL: Component 5 — Instructor header found but only {len(instructor_paras)} entries")
            total_score += 0.05
        else:
            print("FAIL: Component 5 — No instructor section found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Footer with address and phone (0.15 pts)
    try:
        footer_text = ''
        for section in doc.sections:
            if section.footer and section.footer.paragraphs:
                footer_text = ' '.join(p.text for p in section.footer.paragraphs)

        has_address = bool(re.search(r'\d+\s+\w+.*\w{2}\s+\d{5}', footer_text))  # street + state + zip pattern
        has_phone = bool(re.search(r'\(?\d{3}\)?[\s\-\.]\d{3}[\s\-\.]\d{4}', footer_text))
        has_gym_name = 'peak performance' in footer_text.lower()

        if (has_address or has_gym_name) and has_phone:
            print(f"PASS: Component 6 — Footer has address/gym name and phone: '{footer_text[:80]}' (0.15 pts)")
            total_score += 0.15
        elif has_address or has_phone or has_gym_name:
            print(f"PARTIAL: Component 6 — Footer partially complete: address={has_address}, phone={has_phone}, gym={has_gym_name}")
            total_score += 0.07
        else:
            print(f"FAIL: Component 6 — Footer empty or missing address/phone: '{footer_text[:80]}'")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice changes
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
