"""
Reward Script: Create custom table appearance in custom_style.docx
Task ID: writer_tbl_062
Domain: libreoffice_writer

Scoring:
  Component 1: Outer border is 1.5pt (sz=12) solid black on all 4 sides   (0.30 pts)
  Component 2: Inner horizontal borders are 0.5pt (sz=4) solid gray on all inner rows  (0.20 pts)
  Component 3: Inner vertical borders are absent/nil (no inner vertical borders)        (0.15 pts)
  Component 4: Header row cells have dark gray background (~RGB 64,64,64 = #404040)     (0.20 pts)
  Component 5: Header row text is white and bold, data rows have no special formatting  (0.15 pts)
  Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_062'
FILE_PATH = os.path.join(WORKDIR, 'Desktop', 'custom_style.docx')

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def get_attr(elem, attr_name):
    return elem.get('{%s}%s' % (W_NS, attr_name))

def get_cell_border(cell, border_name):
    """Returns (sz, val, color) for a given border of a cell, or (None, None, None)."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return (None, None, None)
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return (None, None, None)
    border = tcBorders.find(qn('w:%s' % border_name))
    if border is None:
        return (None, None, None)
    return (get_attr(border, 'sz'), get_attr(border, 'val'), get_attr(border, 'color'))

def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print('CRITICAL: Cannot load file %s: %s' % (file_path, e))
        print('REWARD: 0.0')
        return 0.0

    if not doc.tables:
        print('CRITICAL: No tables found in document')
        print('REWARD: 0.0')
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    if num_rows < 4 or num_cols < 4:
        print('CRITICAL: Table does not have expected 4x4 dimensions (found %dx%d)' % (num_rows, num_cols))
        print('REWARD: 0.0')
        return 0.0

    # ---------------------------------------------------------------------------
    # Component 1: Outer border is 1.5pt (sz=12) solid black on all 4 sides (0.30 pts)
    #   - Top of row 0 cells: sz=12, val=single, color=000000
    #   - Bottom of last row cells: sz=12, val=single, color=000000
    #   - Left of col 0 cells: sz=12, val=single, color=000000
    #   - Right of last col cells: sz=12, val=single, color=000000
    # ---------------------------------------------------------------------------
    try:
        outer_passes = 0
        outer_total = 0

        # Top border of all header row cells
        for j in range(num_cols):
            sz, val, color = get_cell_border(table.cell(0, j), 'top')
            outer_total += 1
            if sz == '12' and val == 'single' and color == '000000':
                outer_passes += 1
            else:
                print('FAIL: Outer top border at cell [0,%d]: sz=%s, val=%s, color=%s' % (j, sz, val, color))

        # Bottom border of all last row cells
        last_row = num_rows - 1
        for j in range(num_cols):
            sz, val, color = get_cell_border(table.cell(last_row, j), 'bottom')
            outer_total += 1
            if sz == '12' and val == 'single' and color == '000000':
                outer_passes += 1
            else:
                print('FAIL: Outer bottom border at cell [%d,%d]: sz=%s, val=%s, color=%s' % (last_row, j, sz, val, color))

        # Left border of all leftmost column cells
        for i in range(num_rows):
            sz, val, color = get_cell_border(table.cell(i, 0), 'left')
            outer_total += 1
            if sz == '12' and val == 'single' and color == '000000':
                outer_passes += 1
            else:
                print('FAIL: Outer left border at cell [%d,0]: sz=%s, val=%s, color=%s' % (i, sz, val, color))

        # Right border of all rightmost column cells
        last_col = num_cols - 1
        for i in range(num_rows):
            sz, val, color = get_cell_border(table.cell(i, last_col), 'right')
            outer_total += 1
            if sz == '12' and val == 'single' and color == '000000':
                outer_passes += 1
            else:
                print('FAIL: Outer right border at cell [%d,%d]: sz=%s, val=%s, color=%s' % (i, last_col, sz, val, color))

        if outer_passes == outer_total:
            print('PASS: Component 1 — All outer borders are 1.5pt solid black (%d/%d checks) (0.30 pts)' % (outer_passes, outer_total))
            total_score += 0.30
        elif outer_passes >= outer_total * 0.75:
            print('PARTIAL: Component 1 — Outer borders mostly correct (%d/%d checks) (0.15 pts)' % (outer_passes, outer_total))
            total_score += 0.15
        else:
            print('FAIL: Component 1 — Outer borders not 1.5pt solid black (%d/%d checks)' % (outer_passes, outer_total))

    except Exception as e:
        print('ERROR: Component 1 — %s' % e)

    # ---------------------------------------------------------------------------
    # Component 2: Inner horizontal borders are 0.5pt (sz=4) solid gray (0.20 pts)
    #   - All cells between rows (bottom of row i for i=0..N-2, top of row i for i=1..N-1)
    #   - These should be sz=4, val=single, color=808080
    # ---------------------------------------------------------------------------
    try:
        inner_h_passes = 0
        inner_h_total = 0

        # Inner horizontal = border between rows:
        # For each row that is not the last row: check 'bottom' borders of those cells
        # (which correspond to horizontal lines between rows)
        # Rows 0 to num_rows-2: their bottom borders are inner horizontal (except for outer bottom which is last row)
        for i in range(num_rows - 1):
            for j in range(num_cols):
                sz, val, color = get_cell_border(table.cell(i, j), 'bottom')
                inner_h_total += 1
                if sz == '4' and val == 'single' and color == '808080':
                    inner_h_passes += 1
                else:
                    print('FAIL: Inner H border at cell [%d,%d] bottom: sz=%s, val=%s, color=%s' % (i, j, sz, val, color))

        if inner_h_passes == inner_h_total:
            print('PASS: Component 2 — All inner horizontal borders are 0.5pt solid gray (%d/%d checks) (0.20 pts)' % (inner_h_passes, inner_h_total))
            total_score += 0.20
        elif inner_h_passes >= inner_h_total * 0.75:
            print('PARTIAL: Component 2 — Inner horizontal borders mostly correct (%d/%d checks) (0.10 pts)' % (inner_h_passes, inner_h_total))
            total_score += 0.10
        else:
            print('FAIL: Component 2 — Inner horizontal borders not 0.5pt solid gray (%d/%d checks)' % (inner_h_passes, inner_h_total))

    except Exception as e:
        print('ERROR: Component 2 — %s' % e)

    # ---------------------------------------------------------------------------
    # Component 3: Inner vertical borders are explicitly suppressed (nil) (0.15 pts)
    #   - In the golden file, the task explicitly sets inner vertical borders to nil.
    #   - This is detected by the presence of a right/left border element with val='nil'.
    #   - The initial file has NO tcBorders at all, which means "not configured" (not suppressed).
    #   - This component only passes when inner vertical cells have EXPLICIT nil settings,
    #     which distinguishes task-configured suppression from the pre-task unconfigured state.
    # ---------------------------------------------------------------------------
    try:
        inner_v_nil_passes = 0
        inner_v_nil_total = 0

        # Check that right border is EXPLICITLY nil for all non-rightmost cells
        # (the element exists and has val='nil' — not just absent/None)
        for i in range(num_rows):
            for j in range(num_cols - 1):
                tc = table.cell(i, j)._tc
                tcPr = tc.find(qn('w:tcPr'))
                tcBorders = tcPr.find(qn('w:tcBorders')) if tcPr is not None else None
                right_elem = tcBorders.find(qn('w:right')) if tcBorders is not None else None
                inner_v_nil_total += 1
                if right_elem is not None and get_attr(right_elem, 'val') == 'nil':
                    inner_v_nil_passes += 1
                else:
                    val_found = get_attr(right_elem, 'val') if right_elem is not None else 'element_missing'
                    print('FAIL: Inner vertical border not explicitly nil at cell [%d,%d] right: val=%s' % (i, j, val_found))

        # Check that left border is EXPLICITLY nil for all non-leftmost cells
        for i in range(num_rows):
            for j in range(1, num_cols):
                tc = table.cell(i, j)._tc
                tcPr = tc.find(qn('w:tcPr'))
                tcBorders = tcPr.find(qn('w:tcBorders')) if tcPr is not None else None
                left_elem = tcBorders.find(qn('w:left')) if tcBorders is not None else None
                inner_v_nil_total += 1
                if left_elem is not None and get_attr(left_elem, 'val') == 'nil':
                    inner_v_nil_passes += 1
                else:
                    val_found = get_attr(left_elem, 'val') if left_elem is not None else 'element_missing'
                    print('FAIL: Inner vertical border not explicitly nil at cell [%d,%d] left: val=%s' % (i, j, val_found))

        if inner_v_nil_passes == inner_v_nil_total:
            print('PASS: Component 3 — All inner vertical borders explicitly nil (%d/%d checks) (0.15 pts)' % (inner_v_nil_passes, inner_v_nil_total))
            total_score += 0.15
        elif inner_v_nil_passes >= inner_v_nil_total * 0.75:
            print('PARTIAL: Component 3 — Inner vertical borders mostly explicit nil (%d/%d checks) (0.08 pts)' % (inner_v_nil_passes, inner_v_nil_total))
            total_score += 0.08
        else:
            print('FAIL: Component 3 — Inner vertical borders not explicitly suppressed (%d/%d checks)' % (inner_v_nil_passes, inner_v_nil_total))

    except Exception as e:
        print('ERROR: Component 3 — %s' % e)

    # ---------------------------------------------------------------------------
    # Component 4: Header row cells have dark gray background (#404040) (0.20 pts)
    # ---------------------------------------------------------------------------
    try:
        header_bg_passes = 0
        header_bg_total = num_cols

        for j in range(num_cols):
            cell = table.cell(0, j)
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            shd = tcPr.find(qn('w:shd')) if tcPr is not None else None
            fill = get_attr(shd, 'fill') if shd is not None else None
            if fill is not None and fill.upper() == '404040':
                header_bg_passes += 1
            else:
                print('FAIL: Header cell [0,%d] background: expected 404040, got %s' % (j, fill))

        # Also check data rows have no special background
        data_rows_bad = 0
        for i in range(1, num_rows):
            for j in range(num_cols):
                cell = table.cell(i, j)
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                shd = tcPr.find(qn('w:shd')) if tcPr is not None else None
                fill = get_attr(shd, 'fill') if shd is not None else None
                # None or auto/white are acceptable for data rows
                if fill is not None and fill.upper() not in ('FFFFFF', 'AUTO', 'NONE', 'F8F8F8') and fill != 'None':
                    data_rows_bad += 1
                    print('FAIL: Data cell [%d,%d] has unexpected background: %s' % (i, j, fill))
        data_rows_ok = (data_rows_bad == 0)

        if header_bg_passes == header_bg_total and data_rows_ok:
            print('PASS: Component 4 — Header row has dark gray (#404040) background, data rows clean (%d/%d cells) (0.20 pts)' % (header_bg_passes, header_bg_total))
            total_score += 0.20
        elif header_bg_passes == header_bg_total:
            print('PARTIAL: Component 4 — Header row has correct background but data rows may have background (0.10 pts)')
            total_score += 0.10
        elif header_bg_passes >= num_cols * 0.5:
            print('PARTIAL: Component 4 — Header background partial (%d/%d cells) (0.10 pts)' % (header_bg_passes, header_bg_total))
            total_score += 0.10
        else:
            print('FAIL: Component 4 — Header background not dark gray (%d/%d cells)' % (header_bg_passes, header_bg_total))

    except Exception as e:
        print('ERROR: Component 4 — %s' % e)

    # ---------------------------------------------------------------------------
    # Component 5: Header row text is white (#FFFFFF) and bold; data rows have default formatting (0.15 pts)
    # ---------------------------------------------------------------------------
    try:
        header_text_passes = 0
        header_text_total = num_cols

        for j in range(num_cols):
            cell = table.cell(0, j)
            cell_ok = False
            for para in cell.paragraphs:
                for run in para.runs:
                    is_bold = run.bold is True or run.font.bold is True
                    run_color = run.font.color.rgb if run.font.color and run.font.color.type is not None else None
                    color_ok = (run_color is not None and str(run_color).upper() == 'FFFFFF')
                    if is_bold and color_ok:
                        cell_ok = (is_bold and color_ok)
            if cell_ok:
                header_text_passes += 1
            else:
                # Check if bold or color individually
                for para in cell.paragraphs:
                    for run in para.runs:
                        is_bold = run.bold is True or run.font.bold is True
                        run_color = run.font.color.rgb if run.font.color and run.font.color.type is not None else None
                        print('FAIL: Header cell [0,%d] run: bold=%s, color=%s' % (j, is_bold, run_color))

        if header_text_passes == header_text_total:
            print('PASS: Component 5 — Header row text is white and bold (%d/%d cells) (0.15 pts)' % (header_text_passes, header_text_total))
            total_score += 0.15
        elif header_text_passes >= num_cols * 0.5:
            print('PARTIAL: Component 5 — Header text partial (%d/%d cells) (0.08 pts)' % (header_text_passes, header_text_total))
            total_score += 0.08
        else:
            print('FAIL: Component 5 — Header text not white/bold (%d/%d cells)' % (header_text_passes, header_text_total))

    except Exception as e:
        print('ERROR: Component 5 — %s' % e)

    # ---------------------------------------------------------------------------
    # Final score
    # ---------------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print('\nScore: %.2f/1.0' % total_score)
    print('REWARD: %.1f' % final_score)
    return final_score


if not os.path.exists(FILE_PATH):
    print('File not found: %s' % FILE_PATH)
    print('REWARD: 0.0')
else:
    verify_task(FILE_PATH)
