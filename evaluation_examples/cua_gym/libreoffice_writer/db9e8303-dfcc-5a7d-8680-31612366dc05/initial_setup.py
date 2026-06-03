"""
Initial Setup: Create a docx file with a 4x4 table with default formatting
Task ID: writer_tbl_062
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user/Desktop'
TASK_ID = 'custom_style'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure the Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a brief introductory paragraph
    doc.add_paragraph("Product Catalog")

    # Create a 4-row, 4-column table with default (no custom) formatting
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # Row 1 (header): 'Code' | 'Name' | 'Category' | 'Price'
    header_data = ['Code', 'Name', 'Category', 'Price']
    for col_idx, text in enumerate(header_data):
        cell = table.cell(0, col_idx)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        # Default formatting only — no bold, no background color

    # Row 2
    row2_data = ['P001', 'Widget', 'Parts', '$12.99']
    for col_idx, text in enumerate(row2_data):
        table.cell(1, col_idx).text = text

    # Row 3
    row3_data = ['P002', 'Gizmo', 'Parts', '$8.50']
    for col_idx, text in enumerate(row3_data):
        table.cell(2, col_idx).text = text

    # Row 4
    row4_data = ['P003', 'Thingamajig', 'Tools', '$24.00']
    for col_idx, text in enumerate(row4_data):
        table.cell(3, col_idx).text = text

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
