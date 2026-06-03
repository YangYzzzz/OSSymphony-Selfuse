"""
Initial Setup: Legal contract document with terms flagged as spelling errors
Task ID: writer_fp_035
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_035'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Between Meridian Holdings, LLC and Pacific Coast Enterprises, Inc.')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()  # spacer

    # Section 1: Definitions
    h1 = doc.add_heading('ARTICLE I: DEFINITIONS AND RECITALS', level=1)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    run = p1.add_run(
        'This Professional Services Agreement ("Agreement") is entered into as of March 15, 2025, '
        'by and between Meridian Holdings, LLC ("Service Provider") and Pacific Coast Enterprises, Inc. '
        '("Client"). The parties acknowledge that matters of indemnification and subrogation shall be '
        'governed by the provisions set forth herein. Any dispute involving estoppel or claims by a '
        'tortfeasor shall be resolved through the arbitration procedures described in Article VII.'
    )
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run = p2.add_run(
        'For the purposes of this Agreement, the following terms shall have the meanings ascribed to them: '
        '"Adjudication" refers to the formal legal process by which a court resolves a dispute. '
        '"Arbitrable" matters are those subject to resolution through binding arbitration. '
        '"Counterclaim" means any claim brought by the respondent against the claimant in the same proceeding.'
    )
    run.font.size = Pt(11)

    # Section 2: Scope of Services
    doc.add_heading('ARTICLE II: SCOPE OF SERVICES AND OBLIGATIONS', level=1)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    run = p3.add_run(
        'The Service Provider shall provide professional consulting services as outlined in Exhibit A. '
        'In the event of a deposition request related to this Agreement, both parties agree to cooperate '
        'fully. The fiduciary duties of each party are outlined in Section 2.3 below. Either party may '
        'seek a writ of habeas corpus if personnel are unlawfully detained in connection with services '
        'rendered under this Agreement.'
    )
    run.font.size = Pt(11)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    run = p4.add_run(
        'The Client acknowledges that injunctive relief may be sought by the Service Provider in cases '
        'of material breach. Matters of jurisprudence applicable to this Agreement shall be determined '
        'by the laws of the State of California. Any lien placed on assets related to services under '
        'this Agreement must be disclosed within thirty (30) calendar days.'
    )
    run.font.size = Pt(11)

    # Section 3: Liability and Indemnification
    doc.add_heading('ARTICLE III: LIABILITY, INDEMNIFICATION, AND REMEDIES', level=1)

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(6)
    run = p5.add_run(
        'In cases of malfeasance by either party, the aggrieved party shall be entitled to pursue all '
        'available legal remedies. Claims of negligence must be supported by documented evidence of '
        'breach of the applicable standard of care. The obligee under any performance bond shall notify '
        'the surety within fifteen (15) business days of any default.'
    )
    run.font.size = Pt(11)

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(6)
    run = p6.add_run(
        'The plaintiff in any action arising under this Agreement must first attempt resolution through '
        'the mediation process described in Article VI. A quorum of the advisory board, consisting of '
        'at least five (5) members, must approve any settlement exceeding $500,000. Either party may '
        'seek rescission of this Agreement if material misrepresentation is proven.'
    )
    run.font.size = Pt(11)

    # Section 4: Dispute Resolution
    doc.add_heading('ARTICLE IV: DISPUTE RESOLUTION AND DEMURRER PROCEDURES', level=1)

    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(6)
    run = p7.add_run(
        'Any party may file a demurrer to challenge the legal sufficiency of claims brought under this '
        'Agreement. The filing of a demurrer shall not stay the obligation of the parties to continue '
        'performing under the terms of this Agreement. Counterclaims must be filed within sixty (60) '
        'days of the initial claim. The adjudication of all disputes shall be conducted in accordance '
        'with the Commercial Arbitration Rules of the American Arbitration Association.'
    )
    run.font.size = Pt(11)

    p8 = doc.add_paragraph()
    p8.paragraph_format.space_after = Pt(6)
    run = p8.add_run(
        'Subrogation rights under this Agreement shall transfer automatically upon payment of any '
        'indemnification claim. The tortfeasor in any third-party action may be joined to proceedings '
        'initiated under this Article. Questions of estoppel shall be determined by the arbitrator '
        'based on the conduct of the parties and the principles of equity and jurisprudence.'
    )
    run.font.size = Pt(11)

    # Section 5: Miscellaneous
    doc.add_heading('ARTICLE V: MISCELLANEOUS PROVISIONS', level=1)

    p9 = doc.add_paragraph()
    p9.paragraph_format.space_after = Pt(6)
    run = p9.add_run(
        'This Agreement constitutes the entire understanding between the parties with respect to the '
        'subject matter hereof. The fiduciary relationship established herein shall survive the '
        'termination of this Agreement for a period of two (2) years. Any arbitrable dispute not '
        'resolved within ninety (90) days of the initial filing shall proceed to binding arbitration. '
        'The negligence standard applicable to services rendered under this Agreement shall be that '
        'of a reasonably prudent professional in the same field.'
    )
    run.font.size = Pt(11)

    p10 = doc.add_paragraph()
    p10.paragraph_format.space_after = Pt(6)
    run = p10.add_run(
        'All notices required under this Agreement shall be delivered to the addresses set forth in '
        'Exhibit B. The obligee and the plaintiff in any enforcement action shall bear their own costs '
        'and attorney fees unless otherwise ordered by the tribunal. A quorum of the oversight committee '
        'is required for any amendment to the terms of rescission or the scope of injunctive remedies '
        'available under this Agreement.'
    )
    run.font.size = Pt(11)

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph()
    sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sig.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Signature lines
    for party, name, title_text in [
        ('MERIDIAN HOLDINGS, LLC', 'Alexandra Thornton', 'Chief Executive Officer'),
        ('PACIFIC COAST ENTERPRISES, INC.', 'David Chen-Ramirez', 'Managing Director')
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(party)
        run.bold = True
        run.font.size = Pt(11)

        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(24)
        run = p_line.add_run('_' * 40)
        run.font.size = Pt(11)

        p_name = doc.add_paragraph()
        run = p_name.add_run(f'Name: {name}')
        run.font.size = Pt(11)

        p_title = doc.add_paragraph()
        run = p_title.add_run(f'Title: {title_text}')
        run.font.size = Pt(11)

        p_date = doc.add_paragraph()
        run = p_date.add_run('Date: March 15, 2025')
        run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure auto spell-check is enabled (set in LibreOffice preferences)
    # Enable auto spell checking via registrymodifications.xcu
    reg_file = '/home/user/.config/libreoffice/4/user/registrymodifications.xcu'
    import xml.etree.ElementTree as ET

    tree = ET.parse(reg_file)
    root = tree.getroot()

    # Add auto spell check setting
    ns = {'oor': 'http://openoffice.org/2001/registry'}

    # We need to add the auto-spellcheck entry if not present
    # The entry is: /org.openoffice.Office.Linguistic/SpellChecking/IsAutoSpellChecking = true
    found = False
    for item in root:
        path = item.get('{http://openoffice.org/2001/registry}path', '')
        if 'IsAutoSpellChecking' in path or 'SpellChecking' in path:
            found = True
            break

    if not found:
        # Add the element
        new_item = ET.SubElement(root, 'item')
        new_item.set('{http://openoffice.org/2001/registry}path',
                     '/org.openoffice.Office.Linguistic/SpellChecking')
        prop = ET.SubElement(new_item, 'prop')
        prop.set('{http://openoffice.org/2001/registry}name', 'IsSpellAuto')
        prop.set('{http://openoffice.org/2001/registry}op', 'fuse')
        value = ET.SubElement(prop, 'value')
        value.text = 'true'

    tree.write(reg_file, xml_declaration=True, encoding='UTF-8')
    print('Auto spell-check enabled')

    # No custom dictionaries should exist - verify wordbook dir is empty/absent
    import shutil
    wordbook_dir = '/home/user/.config/libreoffice/4/user/wordbook'
    if os.path.exists(wordbook_dir):
        shutil.rmtree(wordbook_dir)
        print('Removed existing wordbook directory to ensure clean state')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
