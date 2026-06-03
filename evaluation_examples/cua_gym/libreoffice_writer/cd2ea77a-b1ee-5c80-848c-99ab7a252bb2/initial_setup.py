"""
Initial Setup: Legal contract with placeholder for cross-reference in indemnification clause
Task ID: writer_legal_022
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_022'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bookmark(paragraph, bookmark_name, run_with_text):
    """Add a bookmark wrapping a run in a paragraph using XML."""
    import random
    bm_id = str(random.randint(1000, 99999))

    # Create bookmarkStart element
    bm_start = paragraph._element.makeelement(
        qn('w:bookmarkStart'),
        {qn('w:id'): bm_id, qn('w:name'): bookmark_name}
    )
    # Create bookmarkEnd element
    bm_end = paragraph._element.makeelement(
        qn('w:bookmarkEnd'),
        {qn('w:id'): bm_id}
    )

    # Insert bookmarkStart before the run, bookmarkEnd after
    run_element = run_with_text._element
    run_element.addprevious(bm_start)
    run_element.addnext(bm_end)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('MASTER SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph()
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = intro.add_run(
        'This Master Services Agreement ("Agreement") is entered into as of March 15, 2025 '
        '("Effective Date"), by and between Meridian Technologies, Inc., a Delaware corporation '
        'with its principal office at 2400 Innovation Drive, Suite 800, San Jose, CA 95134 '
        '("Company"), and Brightpath Consulting Group, LLC, a New York limited liability company '
        'with its principal office at 350 Fifth Avenue, 21st Floor, New York, NY 10118 ("Consultant").'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Section 1: General Provisions ---
    h1 = doc.add_heading('Section 1: General Provisions', level=1)

    # Section 1.1
    doc.add_heading('1.1 Purpose', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(
        'The purpose of this Agreement is to establish the terms and conditions under which '
        'Consultant shall provide professional services to Company. These services include but '
        'are not limited to strategic technology consulting, software architecture review, and '
        'systems integration advisory services.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    # Section 1.2
    doc.add_heading('1.2 Scope of Services', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(
        'Consultant agrees to perform the services described in each Statement of Work ("SOW") '
        'executed by both parties. Each SOW shall reference this Agreement and shall be deemed '
        'incorporated herein. In the event of a conflict between this Agreement and any SOW, '
        'the terms of this Agreement shall control unless the SOW expressly states otherwise.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    # Section 1.3
    doc.add_heading('1.3 Term', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(
        'This Agreement shall commence on the Effective Date and shall continue for a period '
        'of three (3) years unless earlier terminated in accordance with Section 10 hereof. '
        'The Agreement may be renewed for successive one-year periods upon written agreement '
        'of both parties executed no later than thirty (30) days prior to the expiration of the '
        'then-current term.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    # Section 1.4
    doc.add_heading('1.4 Governing Law', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(
        'This Agreement shall be governed by and construed in accordance with the laws of the '
        'State of Delaware, without regard to its conflict of laws principles. Any disputes '
        'arising under or in connection with this Agreement shall be subject to the exclusive '
        'jurisdiction of the state and federal courts located in New Castle County, Delaware.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    # Section 1.5 - Definitions (WITH BOOKMARK)
    h_def = doc.add_heading('1.5 Definitions', level=2)

    p_def = doc.add_paragraph()
    p_def.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r_def = p_def.add_run('Section 1.5')
    r_def.font.size = Pt(11)
    r_def.font.name = 'Calibri'
    r_def.bold = True
    # Add bookmark on this run
    add_bookmark(p_def, 'definitions_section', r_def)

    # Add the definitions text as a continuation
    r2 = p_def.add_run(
        ' - For the purposes of this Agreement, the following terms shall have the meanings '
        'set forth below:'
    )
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'

    definitions = [
        ('"Confidential Information"', 'means any non-public information disclosed by either party '
         'to the other party, whether orally, in writing, or electronically, that is designated as '
         'confidential or that reasonably should be understood to be confidential given the nature '
         'of the information and the circumstances of disclosure.'),
        ('"Deliverables"', 'means all work product, reports, analyses, software, documentation, '
         'and other materials produced by Consultant in the course of performing services under '
         'this Agreement or any SOW.'),
        ('"Intellectual Property"', 'means all patents, copyrights, trademarks, trade secrets, '
         'and other intellectual property rights in and to the Deliverables and any pre-existing '
         'materials incorporated therein.'),
        ('"Losses"', 'means all claims, damages, liabilities, costs, and expenses, including '
         'reasonable attorneys\' fees and court costs, arising from or related to any breach of '
         'this Agreement or any negligent or wrongful act or omission.'),
        ('"Personnel"', 'means the employees, agents, and subcontractors of Consultant who are '
         'assigned to perform services under this Agreement.'),
    ]

    for term, definition in definitions:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        rt = p.add_run(term)
        rt.bold = True
        rt.font.size = Pt(11)
        rt.font.name = 'Calibri'
        rd = p.add_run(' ' + definition)
        rd.font.size = Pt(11)
        rd.font.name = 'Calibri'

    # --- Section 2-7 (abbreviated but realistic) ---
    sections = [
        ('Section 2: Compensation', [
            'Company shall pay Consultant the fees set forth in each SOW. Unless otherwise specified, '
            'all invoices shall be payable within thirty (30) days of receipt. Late payments shall '
            'accrue interest at the rate of 1.5% per month or the maximum rate permitted by law, '
            'whichever is less.',
            'Consultant shall submit detailed invoices on a monthly basis, itemizing the services '
            'performed, hours worked, and expenses incurred during the preceding calendar month. '
            'All expenses exceeding $500.00 must be pre-approved in writing by Company.'
        ]),
        ('Section 3: Confidentiality', [
            'Each party agrees to hold the other party\'s Confidential Information in strict confidence '
            'and not to disclose such information to any third party without the prior written consent '
            'of the disclosing party. This obligation of confidentiality shall survive the termination '
            'of this Agreement for a period of five (5) years.',
            'The obligations set forth in this Section shall not apply to information that: (a) is or '
            'becomes publicly available through no fault of the receiving party; (b) was already known '
            'to the receiving party prior to disclosure; (c) is independently developed by the receiving '
            'party without reference to the disclosing party\'s Confidential Information; or (d) is '
            'required to be disclosed by law or court order.'
        ]),
        ('Section 4: Intellectual Property', [
            'All Deliverables produced under this Agreement shall be considered "work made for hire" '
            'to the extent permitted by applicable law. To the extent any Deliverable does not qualify '
            'as work made for hire, Consultant hereby assigns to Company all right, title, and interest '
            'in and to such Deliverable.',
            'Consultant retains ownership of any pre-existing intellectual property that Consultant '
            'incorporates into the Deliverables, provided that Consultant grants Company a perpetual, '
            'non-exclusive, royalty-free license to use such pre-existing materials as part of the '
            'Deliverables.'
        ]),
        ('Section 5: Representations and Warranties', [
            'Consultant represents and warrants that: (a) it has the legal right and authority to enter '
            'into this Agreement; (b) the services will be performed in a professional and workmanlike '
            'manner consistent with industry standards; (c) the Deliverables will not infringe upon '
            'any third party\'s intellectual property rights; and (d) Consultant will comply with all '
            'applicable laws and regulations in performing the services.'
        ]),
        ('Section 6: Limitation of Liability', [
            'IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER PARTY FOR ANY INDIRECT, INCIDENTAL, '
            'SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF OR RELATED TO THIS AGREEMENT, '
            'REGARDLESS OF WHETHER SUCH DAMAGES ARE BASED ON CONTRACT, TORT, STRICT LIABILITY, OR ANY '
            'OTHER THEORY.',
            'The total aggregate liability of either party under this Agreement shall not exceed the '
            'total amount of fees paid or payable by Company to Consultant during the twelve (12) month '
            'period immediately preceding the event giving rise to such liability.'
        ]),
        ('Section 7: Insurance', [
            'Consultant shall maintain at all times during the term of this Agreement the following '
            'insurance coverage: (a) commercial general liability insurance with minimum limits of '
            '$2,000,000 per occurrence and $5,000,000 in the aggregate; (b) professional liability '
            '(errors and omissions) insurance with minimum limits of $3,000,000 per claim; and '
            '(c) workers\' compensation insurance as required by applicable law.',
        ]),
    ]

    for section_title, paragraphs in sections:
        doc.add_heading(section_title, level=1)
        for text in paragraphs:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            r = p.add_run(text)
            r.font.size = Pt(11)
            r.font.name = 'Calibri'

    # --- Section 8: Indemnification (with placeholder) ---
    doc.add_heading('Section 8: Indemnification', level=1)

    p8a = doc.add_paragraph()
    p8a.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r8a = p8a.add_run(
        'Consultant shall indemnify, defend, and hold harmless Company and its officers, directors, '
        'employees, and agents from and against any and all Losses, as defined in [see definitions], '
        'arising out of or resulting from: (a) any breach of this Agreement by Consultant; (b) any '
        'negligent or wrongful act or omission of Consultant or its Personnel in the performance of '
        'services under this Agreement; or (c) any claim that the Deliverables infringe upon the '
        'intellectual property rights of any third party.'
    )
    r8a.font.size = Pt(11)
    r8a.font.name = 'Calibri'

    p8b = doc.add_paragraph()
    p8b.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r8b = p8b.add_run(
        'Company shall indemnify, defend, and hold harmless Consultant and its officers, members, '
        'employees, and agents from and against any and all Losses arising out of or resulting from: '
        '(a) any breach of this Agreement by Company; or (b) any negligent or wrongful act or omission '
        'of Company in connection with this Agreement.'
    )
    r8b.font.size = Pt(11)
    r8b.font.name = 'Calibri'

    p8c = doc.add_paragraph()
    p8c.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r8c = p8c.add_run(
        'The indemnifying party\'s obligations under this Section are conditioned upon: (i) the '
        'indemnified party providing prompt written notice of any claim; (ii) the indemnified party '
        'granting the indemnifying party sole control of the defense and settlement of such claim; '
        'and (iii) the indemnified party providing reasonable cooperation at the indemnifying party\'s '
        'expense.'
    )
    r8c.font.size = Pt(11)
    r8c.font.name = 'Calibri'

    # --- Section 9: Non-Solicitation ---
    doc.add_heading('Section 9: Non-Solicitation', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(
        'During the term of this Agreement and for a period of twelve (12) months following its '
        'termination, neither party shall, directly or indirectly, solicit, recruit, or hire any '
        'employee or contractor of the other party who was involved in the performance of services '
        'under this Agreement, without the prior written consent of the other party.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    # --- Section 10: Termination ---
    doc.add_heading('Section 10: Termination', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(
        'Either party may terminate this Agreement upon sixty (60) days\' prior written notice to '
        'the other party. Either party may terminate this Agreement immediately upon written notice '
        'if the other party: (a) materially breaches any provision of this Agreement and fails to '
        'cure such breach within thirty (30) days after receipt of written notice; (b) becomes '
        'insolvent or files for bankruptcy protection; or (c) ceases to conduct business in the '
        'normal course.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Calibri'

    # --- Signature Block ---
    doc.add_paragraph()  # spacer
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r_sig = p_sig.add_run('IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.')
    r_sig.font.size = Pt(11)
    r_sig.font.name = 'Calibri'
    r_sig.bold = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
