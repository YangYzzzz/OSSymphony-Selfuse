"""
Initial Setup: Company letterhead document for termination letter task
Task ID: writer_hr_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Company Letterhead ---
    # Company Name
    company_heading = doc.add_paragraph()
    company_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = company_heading.add_run("GlobalTech Industries")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.font.name = "Arial"

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tag_run = tagline.add_run("Innovating Tomorrow, Today")
    tag_run.italic = True
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    tag_run.font.name = "Arial"

    # Address line
    addr = doc.add_paragraph()
    addr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr_run = addr.add_run("2500 Innovation Drive, Suite 400  |  San Jose, CA 95134  |  (408) 555-7200  |  hr@globaltech.com")
    addr_run.font.size = Pt(9)
    addr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    addr_run.font.name = "Arial"

    # Horizontal rule (thin line separator)
    separator = doc.add_paragraph()
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sep_run = separator.add_run("_" * 72)
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # Empty paragraph for spacing
    doc.add_paragraph()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
