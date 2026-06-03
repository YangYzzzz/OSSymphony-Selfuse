"""
Reward Script: Termination letter for Emily Rodriguez
Task ID: writer_hr_038
Domain: libreoffice_writer
Scoring:
  Component 1: Letter addresses Emily Rodriguez with April 30, 2026 last day (0.20)
  Component 2: Final Paycheck section present with details (0.20)
  Component 3: Benefits Continuation / COBRA section present (0.20)
  Component 4: Return of Company Property section with items (0.20)
  Component 5: Exit Interview section present (0.10)
  Component 6: Formal closing with HR Director signature (0.10)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_038'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph text for searching
    all_text = "\n".join(p.text for p in doc.paragraphs)
    all_text_lower = all_text.lower()

    # The initial file has only 5 paragraphs (letterhead only).
    # The golden file has 39 paragraphs with full letter content.
    # We need the letter body to exist beyond the letterhead.

    # Component 1: Letter addresses Emily Rodriguez with April 30, 2026 (0.20 points)
    # This checks that the letter body was added (not present in initial letterhead-only file)
    try:
        has_emily = "emily rodriguez" in all_text_lower
        # Use regex to match "April 30" near "2026" to be robust against encoding
        has_last_day = bool(re.search(r'[Aa]pril\s+30\S?\s*2026', all_text))
        has_termination_context = any(
            kw in all_text_lower for kw in ["terminat", "last day", "employment"]
        )

        if has_emily and has_last_day and has_termination_context:
            print(f"PASS: Component 1 — Emily Rodriguez addressed, April 30, 2026 found, termination context present (0.20 pts)")
            total_score += 0.20
        elif has_emily and has_last_day:
            print(f"PARTIAL: Component 1 — Emily + date found but no termination context (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — emily={has_emily}, last_day={has_last_day}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Final Paycheck section (0.20 points)
    # Must have a heading/mention of final paycheck AND details about payment
    try:
        has_paycheck_heading = False
        has_paycheck_details = False

        for p in doc.paragraphs:
            text_lower = p.text.strip().lower()
            if "final paycheck" in text_lower and len(p.text.strip()) < 50:
                has_paycheck_heading = True
            if any(kw in text_lower for kw in ["compensation", "pay date", "accrued"]):
                has_paycheck_details = True

        if has_paycheck_heading and has_paycheck_details:
            print(f"PASS: Component 2 — Final Paycheck section with details (0.20 pts)")
            total_score += 0.20
        elif has_paycheck_heading or "final paycheck" in all_text_lower:
            print(f"PARTIAL: Component 2 — Final paycheck mentioned but lacking detail (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — No final paycheck section found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Benefits Continuation / COBRA section (0.20 points)
    try:
        has_cobra_heading = False
        has_cobra_details = False

        for p in doc.paragraphs:
            text_lower = p.text.strip().lower()
            if "cobra" in text_lower and len(p.text.strip()) < 60:
                has_cobra_heading = True
            if any(kw in text_lower for kw in ["health", "dental", "vision", "insurance", "coverage", "18 months"]):
                has_cobra_details = True

        if has_cobra_heading and has_cobra_details:
            print(f"PASS: Component 3 — COBRA/Benefits section with details (0.20 pts)")
            total_score += 0.20
        elif "cobra" in all_text_lower or "benefits continuation" in all_text_lower:
            print(f"PARTIAL: Component 3 — COBRA/benefits mentioned but lacking detail (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No COBRA/benefits continuation section found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Return of Company Property section with bullet items (0.20 points)
    try:
        has_property_heading = False
        has_property_items = False
        bullet_count = 0

        for p in doc.paragraphs:
            text_lower = p.text.strip().lower()
            if "return" in text_lower and "property" in text_lower and len(p.text.strip()) < 60:
                has_property_heading = True
            # Count bullet-style list items related to property
            if p.style and "list" in p.style.name.lower():
                if any(kw in text_lower for kw in ["laptop", "phone", "badge", "card", "document", "charger", "access"]):
                    bullet_count += 1

        has_property_items = bullet_count >= 2

        if has_property_heading and has_property_items:
            print(f"PASS: Component 4 — Company property section with {bullet_count} items (0.20 pts)")
            total_score += 0.20
        elif has_property_heading:
            print(f"PARTIAL: Component 4 — Property heading found but insufficient bullet items ({bullet_count}) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — No company property return section found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Exit Interview section (0.10 points)
    try:
        has_exit_heading = False
        has_exit_details = False

        for p in doc.paragraphs:
            text_lower = p.text.strip().lower()
            if "exit interview" in text_lower and len(p.text.strip()) < 50:
                has_exit_heading = True
            if "exit interview" in text_lower and any(
                kw in text_lower for kw in ["scheduled", "date", "time", "conference", "feedback"]
            ):
                has_exit_details = True

        if has_exit_heading and has_exit_details:
            print(f"PASS: Component 5 — Exit interview section with scheduling details (0.10 pts)")
            total_score += 0.10
        elif "exit interview" in all_text_lower:
            print(f"PARTIAL: Component 5 — Exit interview mentioned but no details (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 5 — No exit interview section found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Formal business letter format with HR Director signature (0.10 points)
    # Checks for: closing salutation ("Sincerely") and signer with HR title
    # Initial file has no closing/signature — only letterhead
    try:
        has_sincerely = False
        has_hr_signer = False

        for p in doc.paragraphs:
            text_lower = p.text.strip().lower()
            if "sincerely" in text_lower or "regards" in text_lower:
                has_sincerely = True
            if "human resources" in text_lower or "hr director" in text_lower:
                has_hr_signer = True

        if has_sincerely and has_hr_signer:
            print(f"PASS: Component 6 — Formal closing with HR signature (0.10 pts)")
            total_score += 0.10
        elif has_sincerely:
            print(f"PARTIAL: Component 6 — Closing found but no HR title (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No formal closing found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
