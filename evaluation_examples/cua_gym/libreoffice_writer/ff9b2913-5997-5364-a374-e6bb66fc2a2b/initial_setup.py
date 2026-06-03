"""
Initial Setup: Create a document with 'Report Date: ' followed by a date field
Task ID: writer_frd_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_date_field(paragraph):
    """Insert a DATE field code into a paragraph (w:fldChar + w:instrText pattern)."""
    # Begin field
    run_begin = paragraph.add_run()
    fld_begin = run_begin._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run_begin._element.append(fld_begin)

    # Field instruction
    run_instr = paragraph.add_run()
    instr_text = run_instr._element.makeelement(qn('w:instrText'), {'{http://www.w3.org/XML/1998/namespace}space': 'preserve'})
    instr_text.text = ' DATE \\@"MM/DD/YYYY" '
    run_instr._element.append(instr_text)

    # Separate
    run_sep = paragraph.add_run()
    fld_sep = run_sep._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run_sep._element.append(fld_sep)

    # Display value (cached)
    run_display = paragraph.add_run('04/02/2026')

    # End field
    run_end = paragraph.add_run()
    fld_end = run_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run_end._element.append(fld_end)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Quarterly Performance Report', level=1)

    # --- Subtitle / metadata paragraph ---
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    run_dept = meta.add_run('Department: Engineering & Product Development')
    run_dept.font.size = Pt(11)
    run_dept.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # --- Report Date paragraph with date field ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(12)
    run_label = date_para.add_run('Report Date: ')
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    add_date_field(date_para)

    # --- Horizontal rule ---
    doc.add_paragraph('_' * 60)

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'This quarterly report summarizes the key achievements, challenges, '
        'and strategic initiatives undertaken by the Engineering and Product '
        'Development department during Q1 2026. Overall team performance has '
        'exceeded expectations with a 15% improvement in sprint velocity and '
        'a 22% reduction in critical bug backlog.'
    )

    # --- Key Metrics ---
    doc.add_heading('Key Performance Metrics', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ['Metric', 'Target', 'Actual']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data = [
        ['Sprint Velocity', '42 story points', '48 story points'],
        ['Bug Resolution Rate', '85%', '91%'],
        ['Code Review Turnaround', '< 24 hours', '18 hours avg'],
        ['Test Coverage', '80%', '87%'],
        ['Customer Satisfaction (NPS)', '> 60', '72'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Team Highlights ---
    doc.add_heading('Team Highlights', level=2)
    highlights = [
        'Successfully launched the new microservices architecture for the payment processing module, reducing transaction latency by 35%.',
        'Completed migration of legacy authentication system to OAuth 2.0, improving security posture across all client applications.',
        'Onboarded three new senior engineers (Sarah Chen, Marcus Rivera, and Priya Patel) who have already contributed to core platform features.',
        'Established weekly cross-functional sync meetings with the Design and QA teams, resulting in fewer handoff-related delays.',
    ]
    for h in highlights:
        doc.add_paragraph(h, style='List Bullet')

    # --- Challenges ---
    doc.add_heading('Challenges & Risks', level=2)
    doc.add_paragraph(
        'The primary challenge this quarter was the unexpected increase in '
        'infrastructure costs due to the scaling requirements of the new '
        'real-time analytics pipeline. The team is working with DevOps to '
        'optimize resource allocation and implement auto-scaling policies '
        'that better match actual usage patterns.'
    )

    # --- Next Quarter Goals ---
    doc.add_heading('Q2 2026 Goals', level=2)
    goals = [
        'Complete Phase 2 of the API gateway modernization project.',
        'Achieve 92% automated test coverage across all critical services.',
        'Reduce mean time to recovery (MTTR) from 45 minutes to under 20 minutes.',
        'Launch the internal developer portal with comprehensive API documentation.',
        'Pilot the new CI/CD pipeline with canary deployment capabilities.',
    ]
    for i, g in enumerate(goals, 1):
        doc.add_paragraph(g, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
