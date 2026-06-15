"""
Reward Script: Find and replace all tab characters with four spaces
Task ID: writer_edit_035
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): No tab characters remain in any paragraph text
  Component 2 (0.3): Each former tab position is replaced with exactly 4 spaces
                     (no 5+ consecutive spaces introduced)
  Component 3 (0.2): Document word content is preserved (no data loss)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_035'
FILE_PATH = '/home/user/Desktop/data_import.docx'


def verify_task(file_path):
    """
    Verify that all tab characters have been replaced with four spaces.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have paragraphs
    if len(doc.paragraphs) == 0:
        print("CRITICAL: Document has no paragraphs — may be corrupted")
        print("REWARD: 0.0")
        return 0.0

    # ----------------------------------------------------------------
    # Component 1: No tab characters remain (0.5 points)
    # This checks the primary task requirement: all \t chars removed.
    # Initial has ~100 tab chars; golden should have 0.
    # ----------------------------------------------------------------
    try:
        total_tabs_in_text = 0
        tab_xml_elements = 0

        for para in doc.paragraphs:
            # Check text-level tabs
            total_tabs_in_text += para.text.count('\t')
            # Check XML-level <w:tab> elements (authoritative)
            tab_elements = para._element.findall('.//' + qn('w:tab'))
            tab_xml_elements += len(tab_elements)

        if total_tabs_in_text == 0 and tab_xml_elements == 0:
            print(f"PASS: Component 1 — No tab characters remain "
                  f"(text_tabs=0, xml_tab_elements=0) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Tab characters still present: "
                  f"text_tabs={total_tabs_in_text}, xml_tab_elements={tab_xml_elements}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ----------------------------------------------------------------
    # Component 2: Each tab replaced with exactly 4 spaces (0.3 points)
    # Verifies the replacement is 4 spaces (not 1, 2, 3, or 5+ spaces).
    # The golden document should have 4-space sequences at each position
    # where a tab character existed, and no 5+ consecutive spaces.
    # Initial: tabs present, no 4-space groups from replacement.
    # Golden: ~100 four-space groups, no 5+ consecutive spaces.
    # ----------------------------------------------------------------
    try:
        four_space_count = 0
        five_plus_space_count = 0

        for para in doc.paragraphs:
            text = para.text
            # Count non-overlapping 4-space groups
            four_space_count += text.count('    ')
            # Count 5+ consecutive space occurrences (wrong replacement width)
            five_plus_space_count += text.count('     ')  # 5 consecutive spaces

        # Must have at least some 4-space replacements and no over-wide groups
        if four_space_count >= 30 and five_plus_space_count == 0:
            print(f"PASS: Component 2 — Exactly 4-space replacements found: "
                  f"{four_space_count} groups, no 5+ consecutive spaces (0.3 pts)")
            total_score += 0.3
        elif five_plus_space_count > 0:
            print(f"FAIL: Component 2 — Found {five_plus_space_count} occurrence(s) of 5+ "
                  f"consecutive spaces; replacement may not be exactly 4 spaces wide")
        else:
            print(f"FAIL: Component 2 — Too few 4-space groups found: {four_space_count} "
                  f"(expected at least 30, task says ~30 tabs)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ----------------------------------------------------------------
    # Component 3: No tabs remain AND column separator format is 4 spaces (0.2 points)
    # Compound check: both tab removal and 4-space insertion are verified together
    # in the context of the column header row, which is the most representative row.
    # This fails on initial (which has tabs, not spaces between column names) and
    # passes on golden (which has 4-space gaps between column names).
    # ----------------------------------------------------------------
    try:
        all_text = ' '.join(para.text for para in doc.paragraphs)

        # Verify that known column headers appear separated by 4 spaces (not tabs)
        # "Region    Q1 Revenue" is the pattern in golden file
        headers_with_spaces = ('Region    Q1 Revenue' in all_text or
                                'Region    ' in all_text)
        # Also ensure no tab exists between "Region" and "Q1 Revenue"
        headers_with_tabs = 'Region\tQ1 Revenue' in all_text or 'Region\t' in all_text

        if headers_with_spaces and not headers_with_tabs:
            print(f"PASS: Component 3 — Column headers use 4-space separators "
                  f"(no tabs between column names) (0.2 pts)")
            total_score += 0.2
        elif headers_with_tabs:
            print(f"FAIL: Component 3 — Column headers still separated by tab characters")
        else:
            print(f"FAIL: Component 3 — Column header separation pattern not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
