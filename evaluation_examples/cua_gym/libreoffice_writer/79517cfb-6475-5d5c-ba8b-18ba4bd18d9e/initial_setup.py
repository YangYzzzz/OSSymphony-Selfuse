"""
Initial Setup: Find and replace all tab characters with four spaces in data_import.docx
Task ID: writer_edit_035
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_edit_035'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/data_import.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set document title style
    doc.add_heading('Sales Performance Report - Q1 2025', level=0)

    # Add intro paragraph
    intro = doc.add_paragraph(
        'This report summarizes quarterly sales performance data exported from the company ERP system. '
        'The following tables show regional sales figures, product categories, and team performance metrics.'
    )
    intro.paragraph_format.space_after = Pt(6)

    # --- Section 1: Regional Sales Data (tab-delimited) ---
    doc.add_heading('Regional Sales Data', level=1)

    # Header row
    p = doc.add_paragraph()
    p.add_run('Region\tQ1 Revenue\tTarget\tVariance\tRep Name')
    p.runs[0].bold = True

    # Data rows with tab-delimited columns (~10 rows, each with 4 tabs = 40+ tabs total)
    regional_data = [
        ('North America',   '$1,245,800',  '$1,200,000',  '+$45,800',    'Sarah Chen'),
        ('Europe',          '$987,340',    '$1,000,000',  '-$12,660',    'Marcus Johnson'),
        ('Asia Pacific',    '$1,102,500',  '$1,050,000',  '+$52,500',    'Priya Patel'),
        ('Latin America',   '$543,200',    '$550,000',    '-$6,800',     'Carlos Rivera'),
        ('Middle East',     '$328,750',    '$300,000',    '+$28,750',    'Fatima Al-Rashid'),
        ('Africa',          '$189,600',    '$200,000',    '-$10,400',    'Kwame Asante'),
        ('Canada',          '$412,900',    '$400,000',    '+$12,900',    'Emily Tremblay'),
        ('Australia',       '$298,450',    '$280,000',    '+$18,450',    'James O\'Brien'),
    ]

    for region, revenue, target, variance, rep in regional_data:
        p = doc.add_paragraph()
        p.add_run(f'{region}\t{revenue}\t{target}\t{variance}\t{rep}')

    doc.add_paragraph('')  # blank line

    # --- Section 2: Product Category Breakdown (tab-delimited) ---
    doc.add_heading('Product Category Breakdown', level=1)

    # Header row
    p = doc.add_paragraph()
    p.add_run('Category\tUnits Sold\tAvg Price\tTotal Revenue\tGrowth %')
    p.runs[0].bold = True

    product_data = [
        ('Enterprise Software',  '2,340',   '$4,200',   '$9,828,000',   '+12.4%'),
        ('Cloud Services',       '8,750',   '$890',     '$7,787,500',   '+28.7%'),
        ('Hardware',             '1,205',   '$3,100',   '$3,735,500',   '-3.2%'),
        ('Professional Services','450',     '$8,500',   '$3,825,000',   '+5.8%'),
        ('Maintenance & Support','3,680',   '$1,200',   '$4,416,000',   '+9.1%'),
        ('Training & Cert.',     '890',     '$1,850',   '$1,646,500',   '+15.3%'),
    ]

    for category, units, avg_price, total_rev, growth in product_data:
        p = doc.add_paragraph()
        p.add_run(f'{category}\t{units}\t{avg_price}\t{total_rev}\t{growth}')

    # Page break to second page
    doc.add_page_break()

    # --- Section 3: Team Performance Metrics (tab-delimited) ---
    doc.add_heading('Team Performance Metrics', level=1)

    p = doc.add_paragraph()
    p.add_run('Sales Rep\tDeals Closed\tAvg Deal Size\tWin Rate\tQuota Attain.')
    p.runs[0].bold = True

    team_data = [
        ('Sarah Chen',        '34',  '$36,641',  '68%',  '103.7%'),
        ('Marcus Johnson',    '28',  '$35,262',  '61%',  '98.7%'),
        ('Priya Patel',       '41',  '$26,890',  '74%',  '104.9%'),
        ('Carlos Rivera',     '22',  '$24,691',  '55%',  '98.8%'),
        ('Fatima Al-Rashid',  '19',  '$17,303',  '63%',  '109.6%'),
        ('Kwame Asante',      '15',  '$12,640',  '58%',  '94.8%'),
        ('Emily Tremblay',    '27',  '$15,293',  '67%',  '103.2%'),
        ('James O\'Brien',    '23',  '$12,976',  '65%',  '106.6%'),
    ]

    for rep, deals, avg_deal, win_rate, quota in team_data:
        p = doc.add_paragraph()
        p.add_run(f'{rep}\t{deals}\t{avg_deal}\t{win_rate}\t{quota}')

    doc.add_paragraph('')  # blank line

    # Add notes section
    doc.add_heading('Notes', level=1)
    notes = doc.add_paragraph(
        'Data extracted from Salesforce CRM on 2025-04-01. All figures in USD unless otherwise noted. '
        'Quota attainment calculated against approved Q1 2025 targets set in January 2025. '
        'Contact revenue-ops@company.com for data discrepancies.'
    )
    notes.paragraph_format.space_before = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count tabs to verify
    tab_count = 0
    doc_verify = Document(OUTPUT)
    for para in doc_verify.paragraphs:
        tab_count += para.text.count('\t')
    print(f'Tab characters in document: {tab_count}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
