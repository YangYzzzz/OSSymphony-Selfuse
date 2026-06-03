"""
Reward Script: Insert a right-pointing arrow shape on page 2 of the document
Task ID: writer_obj_020
Domain: libreoffice_writer
Scoring:
  Component 1: A rightArrow shape exists in the document body     — 0.50 pts
  Component 2: Arrow is located on page 2 (after page break)      — 0.30 pts
  Component 3: Arrow dimensions approx 4cm x 2cm (±20% tolerance) — 0.20 pts
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'process_doc'
FILE_PATH = f'{WORKDIR}/{TASK_ID}.docx'

# EMU per cm
EMU_PER_CM = 360000


def find_page_break_paragraph_index(doc):
    """Return the paragraph index of the first manual page break in the document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                br_type = br.attrib.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
                )
                if br_type == 'page':
                    return i
    return None


def find_right_arrow_shapes(doc):
    """
    Return a list of dicts describing all rightArrow drawing shapes found in the document body.
    Each dict has keys: 'para_index', 'cx_emu', 'cy_emu'
    """
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }

    results = []
    for para_idx, para in enumerate(doc.paragraphs):
        # Check if this paragraph contains a <w:drawing> element
        p_xml = para._element.xml
        if '<w:drawing' not in p_xml:
            continue

        # Find all drawing elements inside this paragraph
        drawings = para._element.findall('.//w:drawing', ns)
        for drawing in drawings:
            # Check inline or anchor
            inline = drawing.find('wp:inline', ns)
            anchor = drawing.find('wp:anchor', ns)

            container = inline if inline is not None else anchor
            if container is None:
                continue

            # Get extent (size)
            extent = container.find('wp:extent', ns)
            cx = int(extent.get('cx', 0)) if extent is not None else 0
            cy = int(extent.get('cy', 0)) if extent is not None else 0

            # Check for rightArrow preset geometry
            prstGeom_elements = drawing.findall('.//a:prstGeom', ns)
            for prstGeom in prstGeom_elements:
                prst = prstGeom.get('prst', '')
                if prst == 'rightArrow':
                    results.append({
                        'para_index': para_idx,
                        'cx_emu': cx,
                        'cy_emu': cy,
                    })

    return results


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must exist and be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the page break that separates page 1 and page 2
    page_break_idx = find_page_break_paragraph_index(doc)
    if page_break_idx is None:
        print("WARN: No manual page break found — cannot determine page 2 boundary")
        # Use a fallback: assume page 2 starts around midpoint
        page_break_idx = len(doc.paragraphs) // 2

    print(f"INFO: Page break found at paragraph index {page_break_idx}")

    # Find all rightArrow shapes in the document
    arrow_shapes = find_right_arrow_shapes(doc)
    print(f"INFO: Found {len(arrow_shapes)} rightArrow shape(s) in document")

    # -------------------------------------------------------------------
    # Component 1: A rightArrow shape exists in the document (0.50 pts)
    # -------------------------------------------------------------------
    try:
        if len(arrow_shapes) > 0:
            shape = arrow_shapes[0]
            print(f"PASS: Component 1 — rightArrow shape found at para {shape['para_index']} (0.50 pts)")
            total_score += 0.50
        else:
            print("FAIL: Component 1 — No rightArrow shape found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------
    # Component 2: Arrow is on page 2 (after page break) (0.30 pts)
    # This ONLY scores if Component 1 passed (arrow exists)
    # -------------------------------------------------------------------
    try:
        if len(arrow_shapes) > 0:
            shape = arrow_shapes[0]
            if shape['para_index'] > page_break_idx:
                print(
                    f"PASS: Component 2 — Arrow is on page 2 "
                    f"(para {shape['para_index']} > page_break {page_break_idx}) (0.30 pts)"
                )
                total_score += 0.30
            else:
                print(
                    f"FAIL: Component 2 — Arrow is on page 1, not page 2 "
                    f"(para {shape['para_index']} <= page_break {page_break_idx})"
                )
        else:
            print("FAIL: Component 2 — No arrow found, cannot check page placement")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------
    # Component 3: Arrow dimensions approximately 4cm x 2cm (0.20 pts)
    # Tolerance: ±20% of the expected EMU values
    # Expected: cx=1440000 EMU (4.0 cm), cy=720000 EMU (2.0 cm)
    # -------------------------------------------------------------------
    try:
        if len(arrow_shapes) > 0:
            shape = arrow_shapes[0]
            cx = shape['cx_emu']
            cy = shape['cy_emu']

            # Expected sizes
            expected_cx = 4 * EMU_PER_CM   # 1440000 EMU
            expected_cy = 2 * EMU_PER_CM   # 720000 EMU
            tolerance = 0.20               # 20% tolerance

            cx_ok = abs(cx - expected_cx) <= expected_cx * tolerance
            cy_ok = abs(cy - expected_cy) <= expected_cy * tolerance

            cx_cm = cx / EMU_PER_CM
            cy_cm = cy / EMU_PER_CM

            if cx_ok and cy_ok:
                print(
                    f"PASS: Component 3 — Arrow size is approximately 4cm x 2cm "
                    f"(actual: {cx_cm:.2f}cm x {cy_cm:.2f}cm) (0.20 pts)"
                )
                total_score += 0.20
            else:
                print(
                    f"FAIL: Component 3 — Arrow size mismatch: "
                    f"expected ~4.0cm x ~2.0cm, "
                    f"got {cx_cm:.2f}cm x {cy_cm:.2f}cm "
                    f"(cx_ok={cx_ok}, cy_ok={cy_ok})"
                )
        else:
            print("FAIL: Component 3 — No arrow found, cannot check dimensions")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Entry point
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
