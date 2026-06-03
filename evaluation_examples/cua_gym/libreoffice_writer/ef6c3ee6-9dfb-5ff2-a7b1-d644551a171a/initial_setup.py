"""
Initial Setup: Create process_doc.docx with a 2-page process description
Task ID: writer_obj_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_020'
OUTPUT = f'{WORKDIR}/process_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page 1: Process Overview ---
    # Title
    title = doc.add_heading('Project Onboarding Process', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction paragraph
    intro = doc.add_paragraph(
        'This document outlines the standard onboarding process for new team members '
        'joining the Engineering and Operations departments. Following this process '
        'ensures consistency and efficiency in bringing new employees up to speed.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Section 1
    doc.add_heading('1. Pre-Arrival Preparation', level=1)
    doc.add_paragraph(
        'Before the new team member arrives, the HR department and direct manager '
        'must complete several preparatory steps to ensure a smooth start.'
    )
    doc.add_paragraph('Set up workstation and required hardware', style='List Bullet')
    doc.add_paragraph('Create user accounts for all required systems', style='List Bullet')
    doc.add_paragraph('Prepare welcome package and access credentials', style='List Bullet')
    doc.add_paragraph('Assign an onboarding buddy from the same department', style='List Bullet')
    doc.add_paragraph('Schedule orientation meetings with key stakeholders', style='List Bullet')

    # Section 2
    doc.add_heading('2. Day 1 Activities', level=1)
    doc.add_paragraph(
        'The first day is crucial for making a positive impression and setting '
        'expectations for the new employee. The following activities should be '
        'completed on the first day of employment.'
    )
    doc.add_paragraph(
        'Morning: The manager conducts a department tour and introduces the new '
        'employee to all team members. Office facilities, meeting rooms, and common '
        'areas are pointed out.'
    )
    doc.add_paragraph(
        'Afternoon: HR completes all required paperwork including benefits enrollment, '
        'tax forms, and compliance training acknowledgments. System access is verified '
        'and any technical issues are resolved.'
    )

    # Section 3
    doc.add_heading('3. First Week Schedule', level=1)
    doc.add_paragraph(
        'During the first week, new employees follow a structured learning plan '
        'designed to provide comprehensive exposure to company systems, culture, '
        'and team responsibilities.'
    )

    # Table for first week schedule
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    # Headers
    headers = ['Day', 'Focus Area', 'Responsible Party']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    schedule_data = [
        ['Monday', 'System access and IT orientation', 'IT Department'],
        ['Tuesday', 'Product and service overview', 'Product Team'],
        ['Wednesday', 'Role-specific training begins', 'Direct Manager'],
        ['Thursday', 'Cross-team collaboration overview', 'Team Lead'],
        ['Friday', 'Goals and KPIs discussion', 'Direct Manager'],
    ]
    for i, row_data in enumerate(schedule_data, 1):
        for j, val in enumerate(row_data):
            table.cell(i, j).text = val

    # Page break to start page 2
    doc.add_page_break()

    # --- Page 2: Process Flow and Additional Steps ---
    doc.add_heading('4. Technical Setup Process', level=1)
    doc.add_paragraph(
        'The technical setup process ensures that new team members have all necessary '
        'tools, access, and configurations to perform their job effectively from day one. '
        'This section describes the step-by-step flow for completing technical onboarding.'
    )

    # Step descriptions - realistic process flow content
    doc.add_heading('Step 1: Account Creation', level=2)
    doc.add_paragraph(
        'The IT team creates accounts for all required platforms: email, VPN, '
        'project management tools (Jira, Confluence), version control (GitHub), '
        'and cloud infrastructure access (AWS, GCP). Each account is provisioned '
        'with role-appropriate permissions following the principle of least privilege.'
    )

    doc.add_heading('Step 2: Device Configuration', level=2)
    doc.add_paragraph(
        'The assigned laptop or workstation is configured with the standard software '
        'stack: operating system updates, security patches, development tools, '
        'communication applications, and department-specific software. '
        'Configuration takes approximately 2-3 hours to complete.'
    )

    doc.add_heading('Step 3: Network and VPN Setup', level=2)
    doc.add_paragraph(
        'Remote access credentials are provisioned and VPN client software is '
        'installed and tested. The employee verifies connectivity to all internal '
        'resources. Two-factor authentication is configured for all sensitive systems.'
    )

    # Note about the process flow visualization
    note_para = doc.add_paragraph()
    note_run = note_para.add_run('Process Flow Overview')
    note_run.bold = True
    note_run.font.size = Pt(12)

    doc.add_paragraph(
        'The diagram below illustrates the sequential flow of the technical setup '
        'process. Each step must be completed before proceeding to the next stage. '
        'The IT department coordinator tracks progress through the onboarding checklist.'
    )

    # Additional content to make it clearly page 2 content
    doc.add_heading('5. Completion and Sign-off', level=1)
    doc.add_paragraph(
        'Upon successful completion of all onboarding activities, the new employee '
        'and manager both sign the onboarding completion form. This form is filed '
        'with HR and serves as confirmation that all required steps have been completed.'
    )

    numbered_steps = [
        'Employee confirms access to all required systems',
        'Manager reviews and approves completion checklist',
        'HR archives onboarding documentation',
        'Feedback survey sent to new employee after 30 days',
    ]
    for step in numbered_steps:
        doc.add_paragraph(step, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
