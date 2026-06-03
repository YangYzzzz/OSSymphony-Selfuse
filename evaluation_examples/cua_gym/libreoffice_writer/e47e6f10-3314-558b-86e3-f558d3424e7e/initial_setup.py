"""
Initial Setup: Writer document with no hyphenation configured
Task ID: writer_fs_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_039'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style.font.size = Pt(12)

    # Configure Text Body style
    if 'Text Body' not in [s.name for s in doc.styles]:
        text_body_style = doc.styles.add_style('Text Body', 1)  # 1 = paragraph
    else:
        text_body_style = doc.styles['Text Body']
    text_body_style.font.name = 'Liberation Serif'
    text_body_style.font.size = Pt(12)
    text_body_style.paragraph_format.space_after = Pt(6)
    text_body_style.paragraph_format.line_spacing = 1.15

    # Configure heading styles
    for level in [1, 2, 3]:
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Liberation Sans'
        if level == 1:
            h_style.font.size = Pt(18)
            h_style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            h_style.font.size = Pt(15)
            h_style.font.color.rgb = RGBColor(0x2E, 0x4A, 0x6E)
        else:
            h_style.font.size = Pt(13)
            h_style.font.color.rgb = RGBColor(0x3D, 0x5A, 0x7D)

    # --- Document content ---

    # Title
    title = doc.add_heading('Quarterly Performance Review: Renewable Energy Division', level=1)

    # Introduction section
    doc.add_heading('Executive Summary', level=2)

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'The Renewable Energy Division has demonstrated exceptional performance during '
        'the third quarter of 2025, surpassing revenue projections by approximately '
        'fourteen percent while maintaining operational efficiency targets across all '
        'major business units. This comprehensive review examines the key performance '
        'indicators, strategic milestones, and organizational developments that contributed '
        'to these outstanding results.'
    )

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'Our photovoltaic manufacturing operations achieved unprecedented production '
        'volumes, delivering over 2.3 gigawatts of panel capacity during the reporting '
        'period. The commissioning of our new automated production facility in Phoenix, '
        'Arizona contributed significantly to these improvements, reducing per-unit '
        'manufacturing costs by approximately twenty-two percent compared to the previous '
        'quarter.'
    )

    # Financial Performance section
    doc.add_heading('Financial Performance Overview', level=2)

    doc.add_heading('Revenue Analysis', level=3)

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'Total divisional revenue for the quarter reached $487.3 million, representing '
        'a year-over-year increase of thirty-one percent. The residential solar installation '
        'segment remained our strongest performer, contributing $198.7 million in revenue '
        'with a gross margin of forty-three percent. Commercial installations generated '
        '$156.2 million, while utility-scale projects accounted for the remaining $132.4 '
        'million.'
    )

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'International operations continued their expansion trajectory, with European '
        'markets contributing approximately eighteen percent of total revenue. The recently '
        'established partnerships with distribution networks in Germany, Spain, and the '
        'Netherlands have accelerated market penetration beyond initial projections. Our '
        'manufacturing collaboration with Solartechnische Werke in Munich has enabled '
        'localized production capabilities that significantly reduce logistics overhead.'
    )

    doc.add_heading('Cost Management', level=3)

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'Operating expenses decreased by approximately eight percent despite the significant '
        'revenue growth, demonstrating the effectiveness of our organizational restructuring '
        'initiatives implemented earlier this year. The consolidation of administrative '
        'functions across regional offices eliminated redundancies while preserving customer '
        'service capabilities. Procurement optimization through strategic supplier agreements '
        'contributed an additional $12.8 million in cost savings.'
    )

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'Research and development expenditures totaled $34.5 million for the quarter, '
        'representing approximately seven percent of revenue. These investments focused '
        'primarily on next-generation heterojunction solar cell technology and advanced '
        'energy storage integration systems. Our intellectual property portfolio expanded '
        'with seventeen new patent applications filed during this period.'
    )

    # Operational Highlights
    doc.add_heading('Operational Highlights', level=2)

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'The manufacturing division successfully completed the transition to automated '
        'quality inspection systems across all production facilities. This implementation '
        'reduced defect rates from 2.3 percent to 0.8 percent while simultaneously '
        'increasing throughput by approximately fifteen percent. The workforce development '
        'program trained over 340 technicians in advanced photovoltaic manufacturing '
        'techniques during the quarter.'
    )

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'Customer satisfaction metrics remained consistently high, with our Net Promoter '
        'Score reaching 72 points, the highest in company history. Installation completion '
        'times improved by approximately twenty percent, averaging 4.2 business days for '
        'residential projects compared to 5.3 days in the previous quarter. Our customer '
        'support center maintained an average response time of under three minutes for '
        'technical assistance requests.'
    )

    # Strategic Initiatives
    doc.add_heading('Strategic Initiatives and Future Outlook', level=2)

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'Looking ahead to the fourth quarter, the division is positioned to capitalize '
        'on several significant opportunities. The anticipated regulatory framework changes '
        'in the European Union regarding renewable energy mandates for commercial buildings '
        'are expected to substantially increase demand for our integrated solar solutions. '
        'Additionally, our partnership with Constellation Energy for utility-scale storage '
        'projects represents a potential $250 million revenue opportunity over the next '
        'eighteen months.'
    )

    p = doc.add_paragraph(style='Text Body')
    p.add_run(
        'The organizational transformation initiative continues to progress according to '
        'schedule, with the final phase of regional office consolidation expected to conclude '
        'by the end of November. This restructuring will establish three primary operational '
        'hubs in Phoenix, Charlotte, and Amsterdam, providing comprehensive geographic '
        'coverage while maintaining cost efficiency.'
    )

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
