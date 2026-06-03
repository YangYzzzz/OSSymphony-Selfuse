"""
Reward Script: Book chapter creation in LibreOffice Writer
Task ID: writer_wf_046
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Heading 1 title "Chapter 7: The Industrial Revolution"
  Component 2 (0.20): Four Heading 2 sections with correct names
  Component 3 (0.15): Body text paragraphs under each section (3-4 sentences)
  Component 4 (0.15): 3 footnotes in the document
  Component 5 (0.20): Inventions table (4 cols, 6 rows incl header)
  Component 6 (0.15): Page numbers in footer
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_046'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect heading info
    heading1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
    heading2_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 2']

    # Component 1: Heading 1 with chapter title (0.15 points)
    try:
        found_title = False
        for p in heading1_paras:
            if 'chapter 7' in p.text.lower() and 'industrial revolution' in p.text.lower():
                found_title = True
                break
        if found_title:
            print(f"PASS: Component 1 -- Heading 1 title found: '{heading1_paras[0].text}' (0.15 pts)")
            total_score += 0.15
        else:
            h1_texts = [p.text for p in heading1_paras]
            print(f"FAIL: Component 1 -- Expected Heading 1 with 'Chapter 7' and 'Industrial Revolution', found: {h1_texts}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Four Heading 2 sections with correct names (0.20 points)
    try:
        expected_sections = ['origins in britain', 'key inventions', 'social impact', 'legacy']
        found_sections = [p.text.strip().lower() for p in heading2_paras]
        matches = sum(1 for exp in expected_sections if any(exp in fs for fs in found_sections))
        if matches == 4:
            print(f"PASS: Component 2 -- All 4 Heading 2 sections found (0.20 pts)")
            total_score += 0.20
        elif matches >= 2:
            partial = round(0.20 * matches / 4, 2)
            print(f"PARTIAL: Component 2 -- {matches}/4 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- Expected 4 Heading 2 sections, found: {found_sections}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Body text under each section (0.15 points)
    # Check that there are Normal-style paragraphs with substantial text (3+ sentences ~ 50+ chars)
    try:
        # Group paragraphs by sections (between headings)
        normal_paras = [p for p in doc.paragraphs
                        if p.style.name == 'Normal' and len(p.text.strip()) > 50]
        # We need body text under each of the 4 sections
        # Count sections that have at least 2 Normal paragraphs after them
        sections_with_body = 0
        para_list = list(doc.paragraphs)
        h2_indices = [i for i, p in enumerate(para_list) if p.style.name == 'Heading 2']
        for idx in h2_indices:
            # Count normal paragraphs after this heading until next heading or end
            body_count = 0
            for j in range(idx + 1, len(para_list)):
                if para_list[j].style.name.startswith('Heading'):
                    break
                if para_list[j].style.name == 'Normal' and len(para_list[j].text.strip()) > 50:
                    body_count += 1
            if body_count >= 2:
                sections_with_body += 1

        if sections_with_body >= 4:
            print(f"PASS: Component 3 -- All 4 sections have body text (0.15 pts)")
            total_score += 0.15
        elif sections_with_body >= 2:
            partial = round(0.15 * sections_with_body / 4, 2)
            print(f"PARTIAL: Component 3 -- {sections_with_body}/4 sections have body text ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 -- Only {sections_with_body}/4 sections have adequate body text")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: 3 footnotes in the document (0.15 points)
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        body_elem = doc.element
        footnote_refs = body_elem.findall('.//w:footnoteReference', ns)
        num_footnotes = len(footnote_refs)
        if num_footnotes >= 3:
            print(f"PASS: Component 4 -- {num_footnotes} footnotes found (0.15 pts)")
            total_score += 0.15
        elif num_footnotes >= 1:
            partial = round(0.15 * num_footnotes / 3, 2)
            print(f"PARTIAL: Component 4 -- {num_footnotes}/3 footnotes found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 -- No footnotes found")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: Inventions table with correct structure (0.20 points)
    try:
        tables = doc.tables
        if len(tables) >= 1:
            table = tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            # Check header row
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            expected_headers = ['invention', 'inventor', 'year', 'impact']
            headers_match = all(eh in hc for eh, hc in zip(expected_headers, header_cells))

            sub_score = 0.0
            # 4 columns (0.05)
            if num_cols == 4:
                sub_score += 0.05
                print(f"  Table has 4 columns: PASS")
            else:
                print(f"  Table has {num_cols} columns, expected 4: FAIL")

            # Header row matches (0.05)
            if headers_match:
                sub_score += 0.05
                print(f"  Table headers match: PASS ({header_cells})")
            else:
                print(f"  Table headers mismatch: expected {expected_headers}, got {header_cells}")

            # 6 rows total (header + 5 inventions) (0.05)
            if num_rows >= 6:
                sub_score += 0.05
                print(f"  Table has {num_rows} rows (>= 6): PASS")
            elif num_rows >= 3:
                sub_score += 0.02
                print(f"  Table has {num_rows} rows (partial): PARTIAL")
            else:
                print(f"  Table has {num_rows} rows, expected >= 6: FAIL")

            # Data rows have content (0.05)
            filled_rows = 0
            for ri in range(1, num_rows):
                cells = [cell.text.strip() for cell in table.rows[ri].cells]
                if all(len(c) > 0 for c in cells):
                    filled_rows += 1
            if filled_rows >= 5:
                sub_score += 0.05
                print(f"  {filled_rows} data rows filled: PASS")
            elif filled_rows >= 2:
                sub_score += 0.02
                print(f"  {filled_rows}/5 data rows filled: PARTIAL")
            else:
                print(f"  {filled_rows}/5 data rows filled: FAIL")

            print(f"{'PASS' if sub_score >= 0.18 else 'PARTIAL'}: Component 5 -- Table verification ({sub_score} pts)")
            total_score += sub_score
        else:
            print(f"FAIL: Component 5 -- No tables found in document")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    # Component 6: Page numbers in footer (0.15 points)
    try:
        has_page_number = False
        for sec in doc.sections:
            footer = sec.footer
            for p in footer.paragraphs:
                # Check for PAGE field code in footer
                instrs = p._element.findall(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText'
                )
                for inst in instrs:
                    if inst.text and 'PAGE' in inst.text.upper():
                        has_page_number = True
                        break
                if has_page_number:
                    break
            if has_page_number:
                break
        if has_page_number:
            print(f"PASS: Component 6 -- Page numbers found in footer (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 6 -- No page number field found in footer")
    except Exception as e:
        print(f"ERROR: Component 6 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification (LibreOffice may have unsaved changes)
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
