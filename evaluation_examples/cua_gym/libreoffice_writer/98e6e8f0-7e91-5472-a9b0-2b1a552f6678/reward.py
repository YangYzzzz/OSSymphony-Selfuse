"""
Reward Script: Reference letter in LibreOffice Writer
Task ID: writer_wf_039
Domain: libreoffice_writer
Scoring:
  C1 (0.20) - Company header "Atlas Engineering" in 16pt bold
  C2 (0.10) - Address/contact info block
  C3 (0.10) - Date line present
  C4 (0.10) - Salutation "To Whom It May Concern"
  C5 (0.25) - 4 recommendation paragraphs mentioning David Park
  C6 (0.15) - Paragraphs cover required topics (intro, technical, leadership, recommendation)
  C7 (0.10) - Signature block (sender name, title, contact)
"""

import os
import re
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_039'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_writer", "libreoffice_calc", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def get_non_empty_paragraphs(doc):
    """Return list of paragraphs with non-empty text."""
    return [p for p in doc.paragraphs if p.text.strip()]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    non_empty = get_non_empty_paragraphs(doc)
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Precondition: document must have meaningful content (at least 5 non-empty paragraphs)
    if len(non_empty) < 5:
        print(f"PRECONDITION FAIL: Only {len(non_empty)} non-empty paragraphs (need >= 5)")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Company header "Atlas Engineering" in 16pt bold (0.20 points)
    try:
        header_found = False
        for para in doc.paragraphs[:3]:  # Header should be in first few paragraphs
            if "atlas engineering" in para.text.lower():
                # Check if any run has bold and ~16pt
                for run in para.runs:
                    if run.font.bold and run.font.size:
                        size_pt = run.font.size.pt
                        if abs(size_pt - 16.0) < 1.0:
                            header_found = True
                            break
                    elif run.font.bold and "atlas engineering" in run.text.lower():
                        # Bold found but size may be inherited — partial
                        header_found = True
                        break
                break
        if header_found:
            print(f"PASS: Component 1 — 'Atlas Engineering' header in bold ~16pt (0.20 pts)")
            total_score += 0.20
        else:
            # Check if at least "Atlas Engineering" appears as bold text somewhere at top
            partial = False
            for para in doc.paragraphs[:3]:
                if "atlas engineering" in para.text.lower():
                    for run in para.runs:
                        if run.font.bold:
                            partial = True
                            break
            if partial:
                print(f"PARTIAL: Component 1 — 'Atlas Engineering' bold but size not 16pt (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 1 — 'Atlas Engineering' header not found in bold 16pt")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Address/contact info block (0.10 points)
    try:
        has_address = False
        address_indicators = 0
        for para in doc.paragraphs[:5]:
            text_lower = para.text.lower()
            if any(kw in text_lower for kw in ["boulevard", "street", "avenue", "road", "suite", "drive"]):
                address_indicators += 1
            if any(kw in text_lower for kw in ["phone", "email", "tel"]):
                address_indicators += 1
            if re.search(r'\b[A-Z]{2}\s+\d{5}', para.text):  # state + zip
                address_indicators += 1
        if address_indicators >= 2:
            has_address = True
            print(f"PASS: Component 2 — Address/contact info found ({address_indicators} indicators) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Address/contact info insufficient (found {address_indicators} indicators)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Date line present (0.10 points)
    try:
        date_found = False
        # Look for a date pattern in the document
        date_patterns = [
            r'(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}',
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{4}-\d{2}-\d{2}',
            r'\d{1,2}\s+(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}',
        ]
        for para in doc.paragraphs[:8]:
            for pattern in date_patterns:
                if re.search(pattern, para.text, re.IGNORECASE):
                    date_found = True
                    break
            if date_found:
                break
        if date_found:
            print(f"PASS: Component 3 — Date line found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No date line found in document")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Salutation "To Whom It May Concern" (0.10 points)
    try:
        salutation_found = False
        for para in doc.paragraphs[:10]:
            if "to whom it may concern" in para.text.lower():
                salutation_found = True
                break
        if salutation_found:
            print(f"PASS: Component 4 — 'To Whom It May Concern' salutation found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — 'To Whom It May Concern' not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 4 recommendation paragraphs mentioning David Park (0.25 points)
    try:
        # Body paragraphs: after salutation, before signature
        # Find salutation index and signature index
        sal_idx = None
        sig_idx = None
        for i, para in enumerate(doc.paragraphs):
            if "to whom it may concern" in para.text.lower():
                sal_idx = i
            if para.text.strip().lower() in ["sincerely,", "sincerely", "best regards,",
                                               "best regards", "regards,", "regards",
                                               "respectfully,", "respectfully",
                                               "yours truly,", "yours sincerely,"]:
                sig_idx = i

        body_paras = []
        if sal_idx is not None:
            start = sal_idx + 1
            end = sig_idx if sig_idx is not None else len(doc.paragraphs)
            body_paras = [p for p in doc.paragraphs[start:end] if len(p.text.strip()) > 30]

        david_mentions = sum(1 for p in body_paras if "david" in p.text.lower() or "park" in p.text.lower())

        if len(body_paras) >= 4 and david_mentions >= 2:
            print(f"PASS: Component 5 — {len(body_paras)} body paragraphs, {david_mentions} mention David/Park (0.25 pts)")
            total_score += 0.25
        elif len(body_paras) >= 3 and david_mentions >= 1:
            print(f"PARTIAL: Component 5 — {len(body_paras)} body paragraphs, {david_mentions} David mentions (0.15 pts)")
            total_score += 0.15
        elif len(body_paras) >= 2:
            print(f"PARTIAL: Component 5 — Only {len(body_paras)} body paragraphs (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 5 — Only {len(body_paras)} body paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Paragraphs cover required topics (0.15 points)
    # Topics: introduction/relationship, technical skills, leadership, recommendation
    try:
        topics_found = 0
        body_text = " ".join(p.text.lower() for p in body_paras) if body_paras else all_text

        # Topic 1: Introduction/relationship
        intro_kw = ["recommend", "writing to", "pleasure", "privilege", "worked with", "supervised",
                     "known", "colleague", "introduction"]
        if any(kw in body_text for kw in intro_kw):
            topics_found += 1

        # Topic 2: Technical skills
        tech_kw = ["technical", "skill", "engineering", "software", "system", "architecture",
                    "programming", "expertise", "proficien", "design"]
        if any(kw in body_text for kw in tech_kw):
            topics_found += 1

        # Topic 3: Leadership
        lead_kw = ["leader", "mentor", "team", "manage", "guid", "collaborat", "initiative",
                    "communicat", "inspir"]
        if any(kw in body_text for kw in lead_kw):
            topics_found += 1

        # Topic 4: Strong recommendation
        rec_kw = ["strong", "enthusiast", "recommend", "asset", "endorse", "confident",
                   "hesitat", "wholeheartedly"]
        if any(kw in body_text for kw in rec_kw):
            topics_found += 1

        if topics_found >= 4:
            print(f"PASS: Component 6 — All 4 topics covered (0.15 pts)")
            total_score += 0.15
        elif topics_found >= 3:
            print(f"PARTIAL: Component 6 — {topics_found}/4 topics covered (0.10 pts)")
            total_score += 0.10
        elif topics_found >= 2:
            print(f"PARTIAL: Component 6 — {topics_found}/4 topics covered (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — Only {topics_found}/4 topics covered")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Signature block (sender name, title, contact) (0.10 points)
    try:
        sig_block_score = 0.0
        # Look for sign-off and info after it
        sig_found = False
        sig_parts = 0
        if sig_idx is not None:
            sig_found = True
            # Check paragraphs after the sign-off
            after_sig = [p for p in doc.paragraphs[sig_idx+1:] if p.text.strip()]
            # We expect: name, title, contact
            if len(after_sig) >= 1:
                sig_parts += 1  # Has at least a name
            if len(after_sig) >= 2:
                sig_parts += 1  # Has title
            # Check for contact info (email or phone)
            for p in after_sig:
                if "@" in p.text or re.search(r'\(\d{3}\)\s*\d{3}', p.text) or "phone" in p.text.lower():
                    sig_parts += 1
                    break
        else:
            # Try to find sign-off keywords in last few paragraphs
            for para in doc.paragraphs[-6:]:
                if any(kw in para.text.lower() for kw in ["sincerely", "regards", "respectfully", "truly"]):
                    sig_found = True
                    break

        if sig_found and sig_parts >= 3:
            print(f"PASS: Component 7 — Signature block complete (name+title+contact) (0.10 pts)")
            total_score += 0.10
        elif sig_found and sig_parts >= 2:
            print(f"PARTIAL: Component 7 — Signature block partial ({sig_parts} parts) (0.07 pts)")
            total_score += 0.07
        elif sig_found:
            print(f"PARTIAL: Component 7 — Sign-off found but minimal info (0.03 pts)")
            total_score += 0.03
        else:
            print(f"FAIL: Component 7 — No signature block found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
