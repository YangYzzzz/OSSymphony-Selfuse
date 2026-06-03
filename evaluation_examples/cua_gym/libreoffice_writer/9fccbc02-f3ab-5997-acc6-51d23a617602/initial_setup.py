"""
Initial Setup: Discovery response document with 15 request-response pairs
Task ID: writer_legal_092
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_092'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

# Discovery request topics for realistic legal content
DISCOVERY_TOPICS = [
    ("all documents, communications, and records relating to the breach of contract alleged in Paragraph 12 of the Complaint",
     "Defendant objects to this Request as overly broad and unduly burdensome. Subject to and without waiving said objections, Defendant will produce non-privileged documents in its possession, custody, or control that are responsive to this Request. See attached production log, Bates Nos. DEF-000001 through DEF-000847."),

    ("any and all employment agreements, consulting agreements, or independent contractor agreements between Plaintiff and any current or former employee of Defendant executed between January 1, 2022 and December 31, 2024",
     "Defendant objects on the grounds that this Request seeks information protected by the attorney-client privilege and/or the work product doctrine. Further, Defendant objects that this Request is not reasonably calculated to lead to the discovery of admissible evidence. Without waiving these objections, Defendant states that no such agreements exist."),

    ("all financial statements, balance sheets, income statements, and cash flow statements for Defendant's Western Division for fiscal years 2022, 2023, and 2024",
     "Defendant objects to this Request to the extent it seeks proprietary and confidential business information. Subject to and without waiving said objection, and pursuant to the Protective Order entered on March 15, 2025, Defendant will produce redacted copies of the requested financial statements. See Bates Nos. DEF-000848 through DEF-001203."),

    ("all electronic communications, including but not limited to emails, text messages, and instant messages, between James Harrington and Victoria Chen from June 1, 2023 through September 30, 2023",
     "Defendant objects to this Request as overly broad in scope and unduly burdensome given the volume of communications during the specified period. Without waiving said objection, Defendant will produce responsive, non-privileged email communications. Text messages and instant messages have been preserved but require additional processing time. Defendant requests a 30-day extension to complete this production."),

    ("any and all meeting minutes, agendas, and presentations from the Board of Directors meetings held between March 2023 and November 2024 at which Project Meridian was discussed",
     "Defendant objects to this Request on the grounds that Board meeting minutes contain privileged strategic discussions and attorney-client communications. Defendant further objects that this Request is disproportionate to the needs of this case. Subject to a privilege log to be provided under separate cover, Defendant will produce non-privileged portions of meeting agendas and non-confidential presentation materials."),

    ("all documents relating to Defendant's quality control procedures and testing protocols for the Apex-7 product line, including any modifications made after the incident of August 14, 2023",
     "Defendant objects to the extent this Request seeks post-incident remedial measures, which are inadmissible under Federal Rule of Evidence 407. Without waiving said objection, Defendant will produce quality control procedures and testing protocols in effect as of August 14, 2023. See Bates Nos. DEF-001204 through DEF-001589."),

    ("all insurance policies, including primary and excess policies, that may provide coverage for the claims asserted in this litigation, together with any correspondence with insurers regarding coverage for this matter",
     "Defendant will produce copies of all potentially applicable insurance policies as required by Rule 26(a)(1)(A)(iv). Defendant objects to the production of correspondence with insurers on the grounds that such communications are protected by the attorney-client privilege and/or common interest doctrine. A privilege log will be provided."),

    ("any and all expert reports, analyses, or opinions obtained by Defendant relating to the structural integrity of the Riverside Commerce Center, whether or not Defendant intends to call such experts at trial",
     "Defendant objects to this Request to the extent it seeks materials protected by the work product doctrine and/or consulting expert privilege under Rule 26(b)(4)(D). Defendant will produce expert reports for any expert designated under Rule 26(a)(2) in accordance with the Court's scheduling order."),

    ("all documents and communications relating to any complaints, claims, or lawsuits filed against Defendant by third parties involving the same or similar products or services at issue in this litigation within the past seven years",
     "Defendant objects to this Request as overly broad, unduly burdensome, and not proportional to the needs of this case. The seven-year timeframe is excessive and encompasses products and services materially different from those at issue. Without waiving these objections, Defendant will produce documents relating to complaints involving the Apex-7 product line filed within three years prior to the incident. See Bates Nos. DEF-001590 through DEF-002104."),

    ("all documents relating to Defendant's internal investigation conducted in September and October 2023 following the incident, including interview notes, memoranda, and any corrective action plans",
     "Defendant objects to this Request in its entirety on the grounds that all documents relating to the internal investigation were prepared at the direction of counsel in anticipation of litigation and are therefore protected by the attorney-client privilege and work product doctrine. This objection is supported by the Declaration of General Counsel Margaret Whitfield, filed under seal on February 10, 2025."),

    ("any and all training materials, manuals, and standard operating procedures provided to employees of Defendant's Western Division who were responsible for oversight of the Apex-7 production process",
     "Subject to a general objection regarding the breadth of the term 'any and all,' Defendant will produce training materials and standard operating procedures applicable to Apex-7 production oversight personnel as of August 2023. Defendant notes that certain proprietary manufacturing processes depicted in these materials are subject to the Protective Order. See Bates Nos. DEF-002105 through DEF-002498."),

    ("all documents and communications reflecting or relating to any settlement negotiations or offers of compromise between the parties prior to the filing of this action",
     "Defendant objects to this Request on the grounds that settlement communications are inadmissible and protected under Federal Rule of Evidence 408. This objection is absolute and no documents will be produced in response to this Request."),

    ("any and all documents relating to revenue, profits, and market share for the Apex product line from 2020 through the present, including internal projections, forecasts, and analyst reports",
     "Defendant objects to this Request as overly broad and seeking commercially sensitive information beyond the scope of relevant damages. Without waiving said objection, Defendant will produce annual revenue summaries for the Apex-7 product specifically, for the period 2022 through 2024, subject to the Protective Order. Internal projections and analyst reports are withheld as proprietary. See Bates Nos. DEF-002499 through DEF-002710."),

    ("all communications between Defendant and any governmental or regulatory agency, including but not limited to the Consumer Product Safety Commission and the Environmental Protection Agency, relating to the Apex-7 product from January 2021 to the present",
     "Defendant will produce non-privileged communications with the Consumer Product Safety Commission relating to the Apex-7 product for the period January 2022 through December 2024. Defendant objects to production of EPA communications as irrelevant to the claims and defenses in this action. Defendant further objects to the overbroad timeframe and will limit production to the three-year period preceding the incident. See Bates Nos. DEF-002711 through DEF-003045."),

    ("all documents sufficient to identify each person who participated in, contributed to, or reviewed the decision to discontinue the Apex-7 product line in January 2024, including any analyses or reports considered in making that decision",
     "Defendant objects to this Request to the extent it seeks information protected by the deliberative process privilege and attorney-client privilege. Without waiving said objections, Defendant will identify by name and title the individuals who participated in the formal decision-making process regarding the Apex-7 discontinuation. Analyses prepared by or at the direction of counsel are withheld as privileged. A privilege log entry is included for each withheld document."),
]


def create_initial():
    doc = Document()

    # Title section
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run("PLAINTIFF'S FIRST SET OF REQUESTS FOR PRODUCTION OF DOCUMENTS")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # Case caption
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption_para.add_run("Hartwell Industries, Inc. v. Meridian Technologies Corp.")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    case_info = doc.add_paragraph()
    case_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = case_info.add_run("Case No. 2024-CV-03892 | United States District Court, Northern District of California")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Blank line
    doc.add_paragraph()

    # Preamble
    preamble = doc.add_paragraph()
    run = preamble.add_run(
        "Pursuant to Federal Rules of Civil Procedure 34 and 26(b), Plaintiff Hartwell Industries, Inc. "
        "hereby requests that Defendant Meridian Technologies Corp. produce the following documents and "
        "electronically stored information within thirty (30) days of service of these Requests."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # 15 request-response pairs - all in default style, no bold/italic, no indent
    for i, (request_text, response_text) in enumerate(DISCOVERY_TOPICS, 1):
        # Request paragraph - plain, no bold, no italic
        req_para = doc.add_paragraph()
        run = req_para.add_run(f"REQUEST NO. {i}:")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.bold = False
        run.italic = False

        req_detail = doc.add_paragraph()
        run = req_detail.add_run(f"Please produce {request_text}.")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        # Response paragraph - plain, no indent
        resp_para = doc.add_paragraph()
        run = resp_para.add_run("RESPONSE:")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.bold = False
        run.italic = False

        resp_detail = doc.add_paragraph()
        run = resp_detail.add_run(response_text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

        # Separator blank line between pairs
        if i < 15:
            doc.add_paragraph()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
