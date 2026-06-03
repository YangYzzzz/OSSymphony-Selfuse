"""
Initial Setup: Create a methodology document about correlation analysis
Task ID: writer_acad_060
Domain: libreoffice_writer

Creates a Writer document with academic methodology content describing
correlation analysis, without any equation/formula objects.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_060'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Assessing the Relationship Between Socioeconomic Factors and Academic Performance: A Quantitative Study', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Authors ---
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_a = authors.add_run('Dr. Elena Vasquez, Department of Educational Research\n')
    run_a.font.size = Pt(11)
    run_b = authors.add_run('Prof. Rajesh Nair, Institute for Statistical Sciences\n')
    run_b.font.size = Pt(11)
    run_c = authors.add_run('University of Western Ontario, 2025')
    run_c.font.size = Pt(11)
    run_c.font.italic = True

    # --- Abstract ---
    doc.add_heading('Abstract', level=2)
    abstract = doc.add_paragraph(
        'This study examines the correlation between household income levels and '
        'standardized test scores among secondary school students in three urban '
        'districts. Using data collected from 847 participants over a two-year period '
        '(2023\u20132025), we apply multiple statistical methods to quantify the strength '
        'and direction of the observed relationships. Our findings indicate a moderate '
        'positive correlation between family socioeconomic status and student '
        'performance on mathematics and reading comprehension assessments, while '
        'controlling for variables such as school funding, teacher experience, and '
        'class size.'
    )
    abstract.paragraph_format.space_after = Pt(12)

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        'Understanding the determinants of academic performance has been a central '
        'concern in educational research for decades. Among the many factors studied, '
        'socioeconomic status (SES) has consistently emerged as a significant predictor '
        'of student outcomes (Coleman et al., 1966; Sirin, 2005). However, the '
        'magnitude and nature of this relationship vary across contexts, necessitating '
        'robust quantitative methods to assess it accurately.'
    )
    doc.add_paragraph(
        'This paper contributes to the existing literature by employing correlation '
        'analysis on a large, longitudinal dataset. We focus specifically on the '
        'Pearson product-moment correlation coefficient as our primary measure of '
        'linear association, supplemented by rank-based and partial correlation '
        'techniques to ensure the robustness of our conclusions.'
    )

    # --- Literature Review ---
    doc.add_heading('2. Literature Review', level=2)
    doc.add_paragraph(
        'Prior studies have established a well-documented link between family income '
        'and educational attainment. Reardon (2011) demonstrated that the income '
        'achievement gap among American students has grown by approximately 40% since '
        'the 1970s. Similarly, Chmielewski (2019) conducted a cross-national analysis '
        'of 30 countries and found that socioeconomic gradients in test scores remain '
        'significant even after accounting for school-level resources.'
    )
    doc.add_paragraph(
        'Duncan and Magnuson (2012) argued that early childhood economic conditions '
        'are particularly influential, with effects persisting into secondary education. '
        'Their longitudinal analysis of the Panel Study of Income Dynamics (PSID) '
        'revealed that a $10,000 increase in annual family income during a child\'s '
        'first five years was associated with a 0.15 standard deviation improvement '
        'in later academic scores.'
    )
    doc.add_paragraph(
        'Despite this body of evidence, methodological critiques abound. Harwell and '
        'LeBeau (2010) cautioned that many studies rely on simple bivariate '
        'correlations without adequately controlling for confounding variables, '
        'leading to inflated effect sizes. This motivates our use of both bivariate '
        'and partial correlation approaches.'
    )

    # --- Methodology ---
    doc.add_heading('3. Methodology', level=2)

    doc.add_heading('3.1 Data Collection', level=3)
    doc.add_paragraph(
        'Data were collected from 847 students across 12 secondary schools in the '
        'metropolitan areas of Toronto, Vancouver, and Montreal. Participants were '
        'enrolled in grades 9 through 12 during the 2023\u20132025 academic years. '
        'Household income data were obtained through confidential parent surveys, '
        'while academic performance was measured using provincial standardized '
        'assessments in mathematics and reading comprehension.'
    )

    doc.add_heading('3.2 Variables', level=3)

    # Variables table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Variable', 'Type', 'Description']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    var_data = [
        ['Income', 'Continuous', 'Annual household income (CAD), range $18,500\u2013$245,000'],
        ['MathScore', 'Continuous', 'Provincial math assessment score (0\u2013100)'],
        ['ReadScore', 'Continuous', 'Provincial reading assessment score (0\u2013100)'],
        ['SchoolFunding', 'Continuous', 'Per-pupil expenditure (CAD) at school level'],
        ['TeacherExp', 'Continuous', 'Average years of teaching experience at school'],
    ]
    for r, row_data in enumerate(var_data, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph('')  # spacer

    doc.add_heading('3.3 Statistical Approach', level=3)
    doc.add_paragraph(
        'To assess the strength and direction of the linear relationship between '
        'household income and academic performance, we employ the Pearson '
        'product-moment correlation coefficient. This statistic measures the degree '
        'of linear association between two continuous variables, producing a value '
        'between \u22121 and +1.'
    )
    doc.add_paragraph(
        'The Pearson correlation coefficient r is defined mathematically as follows. '
        'For a sample of n paired observations (x, y), the coefficient captures '
        'the ratio of the covariance of the two variables to the product of their '
        'standard deviations. The formula will be presented below using the equation '
        'editor for clarity:'
    )
    doc.add_paragraph(
        '[Equation: Pearson correlation coefficient r to be inserted here]'
    )
    doc.add_paragraph(
        'A significance test is conducted using the t-distribution with n \u2212 2 '
        'degrees of freedom to determine whether the observed correlation differs '
        'significantly from zero. We adopt a significance level of \u03b1 = 0.05 for '
        'all hypothesis tests.'
    )

    # --- Descriptive Statistics ---
    doc.add_heading('3.4 Descriptive Statistics', level=3)
    doc.add_paragraph(
        'Table 2 presents summary statistics for the primary variables of interest. '
        'All continuous variables were assessed for normality using the '
        'Shapiro\u2013Wilk test prior to correlation analysis.'
    )

    # Descriptive stats table
    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    headers2 = ['Variable', 'Mean', 'Std Dev', 'Min', 'Max']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    stats_data = [
        ['Income (CAD)', '$72,340', '$38,210', '$18,500', '$245,000'],
        ['MathScore', '68.4', '14.7', '22', '99'],
        ['ReadScore', '71.2', '12.3', '31', '98'],
        ['SchoolFunding', '$11,890', '$2,450', '$7,200', '$18,600'],
        ['TeacherExp (yrs)', '12.6', '5.8', '1', '34'],
    ]
    for r, row_data in enumerate(stats_data, 1):
        for c, val in enumerate(row_data):
            cell = table2.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.add_paragraph('')  # spacer

    # --- Preliminary Results mention ---
    doc.add_heading('4. Expected Analysis', level=2)
    doc.add_paragraph(
        'Once the Pearson correlation coefficient formula is applied to our dataset, '
        'we anticipate reporting bivariate correlations between Income and MathScore, '
        'Income and ReadScore, and partial correlations controlling for SchoolFunding '
        'and TeacherExp. The results will be presented in a correlation matrix in '
        'the following section.'
    )

    # --- References ---
    doc.add_heading('References', level=2)
    refs = [
        'Chmielewski, A. K. (2019). The global increase in the socioeconomic achievement gap, 1964 to 2015. American Sociological Review, 84(3), 517\u2013544.',
        'Coleman, J. S., et al. (1966). Equality of Educational Opportunity. U.S. Government Printing Office.',
        'Duncan, G. J., & Magnuson, K. (2012). Socioeconomic status and cognitive functioning. Current Directions in Psychological Science, 21(4), 243\u2013248.',
        'Harwell, M., & LeBeau, B. (2010). Student eligibility for a free lunch as an SES measure in education research. Educational Researcher, 39(2), 120\u2013131.',
        'Reardon, S. F. (2011). The widening academic achievement gap between the rich and the poor. In G. J. Duncan & R. J. Murnane (Eds.), Whither Opportunity? (pp. 91\u2013116). Russell Sage Foundation.',
        'Sirin, S. R. (2005). Socioeconomic status and academic achievement: A meta-analytic review. Review of Educational Research, 75(3), 417\u2013453.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
