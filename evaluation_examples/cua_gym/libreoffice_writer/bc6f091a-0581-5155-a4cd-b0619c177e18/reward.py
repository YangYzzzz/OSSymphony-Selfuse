"""
Reward Script: Create a project charter in LibreOffice Writer
Task ID: writer_wf_062
Domain: libreoffice_writer
Scoring:
  Component 1: Title present with correct text (0.10)
  Component 2: 7 Heading 1 sections with correct names (0.20)
  Component 3: Objectives table with 4 data rows and 3 columns (0.15)
  Component 4: Scope bullets — 5 In-Scope + 3 Out-of-Scope (0.15)
  Component 5: Stakeholders table with 5 data rows and 3 columns (0.15)
  Component 6: 1.5 line spacing on body paragraphs (0.10)
  Component 7: Header contains project ID (0.15)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_062'

EXPECTED_HEADINGS = [
    'Project Overview',
    'Business Case',
    'Objectives and Success Criteria',
    'Scope',
    'Stakeholders',
    'Constraints and Assumptions',
    'Approval',
]

def verify_task(file_path):
    """Verify task completion with progressive scoring. Returns float 0.0-1.0."""
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title present with correct text (0.10 points)
    try:
        title_paras = [p for p in doc.paragraphs
                       if (p.style.name == 'Title' or p.style.name.startswith('Heading'))
                       and 'project charter' in p.text.lower()
                       and 'customer portal redesign' in p.text.lower()]
        if len(title_paras) > 0:
            print(f"PASS: Component 1 — Title found: '{title_paras[0].text}' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Title 'Project Charter - Customer Portal Redesign' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 7 Heading 1 sections with correct names (0.20 points)
    try:
        h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
        h1_texts = [p.text.strip() for p in h1_paras]
        matched = 0
        for expected in EXPECTED_HEADINGS:
            for actual in h1_texts:
                if expected.lower() in actual.lower() or actual.lower() in expected.lower():
                    matched += 1
                    break
        if matched == 7 and len(h1_paras) >= 7:
            print(f"PASS: Component 2 — All 7 Heading 1 sections found: {h1_texts} (0.20 pts)")
            total_score += 0.20
        elif matched >= 5:
            partial = round(0.20 * (matched / 7), 2)
            print(f"PARTIAL: Component 2 — {matched}/7 Heading 1 sections found: {h1_texts} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched}/7 Heading 1 sections found. Got: {h1_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Objectives table with 4 data rows and 3 columns (0.15 points)
    try:
        obj_table = None
        for table in doc.tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'objective' in header_cells or 'metric' in header_cells:
                obj_table = table
                break
        if obj_table is not None:
            num_data_rows = len(obj_table.rows) - 1  # exclude header
            num_cols = len(obj_table.columns)
            if num_data_rows >= 4 and num_cols >= 3:
                print(f"PASS: Component 3 — Objectives table: {num_data_rows} data rows, {num_cols} cols (0.15 pts)")
                total_score += 0.15
            elif num_data_rows >= 2:
                partial = round(0.15 * min(num_data_rows, 4) / 4, 2)
                print(f"PARTIAL: Component 3 — Objectives table: {num_data_rows} data rows (expected 4) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Objectives table too small: {num_data_rows} data rows, {num_cols} cols")
        else:
            print(f"FAIL: Component 3 — Objectives table not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Scope bullets — 5 In-Scope + 3 Out-of-Scope (0.15 points)
    try:
        # Find paragraphs between "Scope" heading and next heading
        scope_start = None
        scope_end = None
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == 'Heading 1' and 'scope' in p.text.lower() and 'out' not in p.text.lower():
                scope_start = i
            elif scope_start is not None and p.style.name == 'Heading 1':
                scope_end = i
                break
        if scope_end is None:
            scope_end = len(doc.paragraphs)

        if scope_start is not None:
            scope_paras = doc.paragraphs[scope_start:scope_end]
            # Find in-scope and out-of-scope sections
            in_scope_bullets = []
            out_scope_bullets = []
            current_section = None
            for p in scope_paras:
                text_lower = p.text.strip().lower()
                if 'in-scope' in text_lower or 'in scope' in text_lower:
                    current_section = 'in'
                elif 'out-of-scope' in text_lower or 'out of scope' in text_lower:
                    current_section = 'out'
                elif p.style.name == 'List Bullet' and p.text.strip():
                    if current_section == 'in':
                        in_scope_bullets.append(p.text.strip())
                    elif current_section == 'out':
                        out_scope_bullets.append(p.text.strip())

            score_4 = 0.0
            if len(in_scope_bullets) >= 5:
                score_4 += 0.09
                print(f"  PASS: In-Scope has {len(in_scope_bullets)} bullets (need 5)")
            elif len(in_scope_bullets) >= 3:
                frac = 0.09 * len(in_scope_bullets) / 5
                score_4 += round(frac, 2)
                print(f"  PARTIAL: In-Scope has {len(in_scope_bullets)}/5 bullets")
            else:
                print(f"  FAIL: In-Scope has {len(in_scope_bullets)} bullets (need 5)")

            if len(out_scope_bullets) >= 3:
                score_4 += 0.06
                print(f"  PASS: Out-of-Scope has {len(out_scope_bullets)} bullets (need 3)")
            elif len(out_scope_bullets) >= 1:
                frac = 0.06 * len(out_scope_bullets) / 3
                score_4 += round(frac, 2)
                print(f"  PARTIAL: Out-of-Scope has {len(out_scope_bullets)}/3 bullets")
            else:
                print(f"  FAIL: Out-of-Scope has {len(out_scope_bullets)} bullets (need 3)")

            if score_4 > 0:
                print(f"PASS: Component 4 — Scope bullets scored ({score_4} pts)")
                total_score += score_4
            else:
                print(f"FAIL: Component 4 — No scope bullets found")
        else:
            print(f"FAIL: Component 4 — Scope section not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Stakeholders table with 5 data rows and 3 columns (0.15 points)
    try:
        stk_table = None
        for table in doc.tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'name' in header_cells and ('role' in header_cells or 'responsibility' in header_cells):
                stk_table = table
                break
        if stk_table is not None:
            num_data_rows = len(stk_table.rows) - 1  # exclude header
            num_cols = len(stk_table.columns)
            if num_data_rows >= 5 and num_cols >= 3:
                print(f"PASS: Component 5 — Stakeholders table: {num_data_rows} data rows, {num_cols} cols (0.15 pts)")
                total_score += 0.15
            elif num_data_rows >= 3:
                partial = round(0.15 * min(num_data_rows, 5) / 5, 2)
                print(f"PARTIAL: Component 5 — Stakeholders table: {num_data_rows} data rows (expected 5) ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — Stakeholders table too small: {num_data_rows} data rows, {num_cols} cols")
        else:
            print(f"FAIL: Component 5 — Stakeholders table not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: 1.5 line spacing on body paragraphs (0.10 points)
    try:
        # Check line spacing on Normal-style paragraphs that have text
        normal_paras = [p for p in doc.paragraphs if p.text.strip() and p.style.name in ('Normal', 'Title', 'List Bullet')]
        if len(normal_paras) == 0:
            print(f"FAIL: Component 6 — No body paragraphs found to check spacing")
        else:
            spacing_ok = 0
            spacing_total = 0
            for p in normal_paras:
                ls = p.paragraph_format.line_spacing
                spacing_total += 1
                if ls is not None and abs(float(ls) - 1.5) < 0.05:
                    spacing_ok += 1
            ratio = spacing_ok / spacing_total if spacing_total > 0 else 0
            if ratio >= 0.8:
                print(f"PASS: Component 6 — 1.5 line spacing on {spacing_ok}/{spacing_total} paragraphs (0.10 pts)")
                total_score += 0.10
            elif ratio >= 0.4:
                partial = round(0.10 * ratio, 2)
                print(f"PARTIAL: Component 6 — 1.5 line spacing on {spacing_ok}/{spacing_total} paragraphs ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — 1.5 line spacing on only {spacing_ok}/{spacing_total} paragraphs")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Header contains project ID (0.15 points)
    try:
        header_text = ''
        for section in doc.sections:
            for hp in section.header.paragraphs:
                header_text += hp.text.strip() + ' '
        header_text = header_text.strip()
        if header_text and ('project id' in header_text.lower() or 'id:' in header_text.lower() or 'cp-' in header_text.lower()):
            print(f"PASS: Component 7 — Header contains project ID: '{header_text}' (0.15 pts)")
            total_score += 0.15
        elif header_text:
            # Has header text but maybe not clearly a project ID; partial credit
            print(f"PARTIAL: Component 7 — Header present but unclear project ID: '{header_text}' (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 7 — Header is empty, expected project ID")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state(domain):
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
