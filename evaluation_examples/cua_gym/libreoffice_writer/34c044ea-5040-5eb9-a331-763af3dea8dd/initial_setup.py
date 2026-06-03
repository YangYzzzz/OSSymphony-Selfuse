"""
Initial Setup: Photo album caption-date layout document (no tabstops)
Task ID: osworld_writer_tabstop_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_tabstop_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # 9 photo album caption lines - no tabs, no tabstops, plain left-aligned text
    captions = [
        "Sunset at the beach - June 2023",
        "Birthday celebration with family - March 2022",
        "Graduation ceremony photo - May 2024",
        "Mountain hike adventure trip - August 2021",
        "Winter holiday gathering at home - December 2023",
        "Spring garden blooming flowers - April 2022",
        "Anniversary dinner restaurant night - February 2024",
        "Summer road trip across countryside - July 2023",
        "New Year fireworks city skyline - January 2024",
    ]

    for caption in captions:
        para = doc.add_paragraph()
        run = para.add_run(caption)
        run.font.size = Pt(12)
        # No tab stops, no alignment changes - plain left-aligned text

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
