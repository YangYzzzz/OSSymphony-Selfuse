"""
Initial Setup: Insert a bookmark at the Troubleshooting section heading
Task ID: writer_tech_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_018'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ===== PAGE 1: Title & Introduction =====
    title = doc.add_heading('CloudSync Platform — Technical Reference Guide', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    run = p.add_run('Version 3.2.1 — April 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacer

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'CloudSync Platform is an enterprise-grade data synchronization solution designed '
        'for organizations managing distributed datasets across multiple cloud providers. '
        'This technical reference guide provides comprehensive documentation for system '
        'administrators, DevOps engineers, and integration specialists responsible for '
        'deploying, configuring, and maintaining CloudSync instances.'
    )
    doc.add_paragraph(
        'The platform supports real-time bidirectional sync between AWS S3, Google Cloud '
        'Storage, and Azure Blob Storage, with configurable conflict resolution policies '
        'and end-to-end encryption using AES-256-GCM. CloudSync handles files up to 5 TB '
        'with automatic chunked transfer and resumable uploads.'
    )

    doc.add_heading('1.1 Scope of This Document', level=2)
    doc.add_paragraph(
        'This guide covers installation procedures, configuration parameters, API usage, '
        'monitoring setup, and troubleshooting steps for CloudSync Platform v3.x. For '
        'migration guides from v2.x, refer to the separate Migration Handbook.'
    )

    doc.add_heading('1.2 Prerequisites', level=2)
    doc.add_paragraph('Before proceeding with installation, ensure the following requirements are met:')
    items = [
        'Linux server (Ubuntu 22.04 LTS or RHEL 9.x) with minimum 8 CPU cores and 32 GB RAM',
        'PostgreSQL 15.x or later for metadata storage',
        'Redis 7.x for caching and job queue management',
        'Valid TLS certificate for the management API endpoint',
        'Network access to target cloud storage endpoints on ports 443 and 8443',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # ===== PAGE 2: Installation & Configuration =====
    doc.add_heading('2. Installation', level=1)
    doc.add_paragraph(
        'CloudSync can be deployed as a standalone binary, a Docker container, or via Helm '
        'chart on Kubernetes. The recommended production deployment uses the Helm chart with '
        'a dedicated namespace and resource quotas.'
    )

    doc.add_heading('2.1 Binary Installation', level=2)
    doc.add_paragraph(
        'Download the latest release archive from the CloudSync releases portal. Extract the '
        'archive and run the installer script:'
    )
    code1 = doc.add_paragraph()
    run_code = code1.add_run(
        '$ tar -xzf cloudsync-3.2.1-linux-amd64.tar.gz\n'
        '$ cd cloudsync-3.2.1\n'
        '$ sudo ./install.sh --prefix=/opt/cloudsync\n'
        '$ sudo systemctl enable --now cloudsync'
    )
    run_code.font.name = 'Courier New'
    run_code.font.size = Pt(9)

    doc.add_heading('2.2 Docker Deployment', level=2)
    doc.add_paragraph(
        'For containerized deployments, pull the official image and mount the configuration '
        'directory as a volume:'
    )
    code2 = doc.add_paragraph()
    run_code2 = code2.add_run(
        '$ docker pull registry.cloudsync.io/cloudsync:3.2.1\n'
        '$ docker run -d --name cloudsync \\\n'
        '    -v /etc/cloudsync:/config \\\n'
        '    -p 8443:8443 -p 9090:9090 \\\n'
        '    registry.cloudsync.io/cloudsync:3.2.1'
    )
    run_code2.font.name = 'Courier New'
    run_code2.font.size = Pt(9)

    doc.add_heading('3. Configuration', level=1)
    doc.add_paragraph(
        'All configuration is managed through a YAML file located at /etc/cloudsync/config.yaml. '
        'The configuration file is validated at startup and any schema violations will prevent '
        'the service from starting.'
    )

    doc.add_heading('3.1 Core Parameters', level=2)

    # Configuration table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    headers = ['Parameter', 'Default', 'Description']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    config_data = [
        ['sync.interval_seconds', '30', 'Polling interval for change detection'],
        ['sync.max_concurrent_jobs', '8', 'Maximum parallel sync operations'],
        ['sync.conflict_policy', 'latest_wins', 'Resolution strategy: latest_wins, source_priority, manual'],
        ['encryption.algorithm', 'AES-256-GCM', 'Encryption cipher for data at rest and in transit'],
        ['storage.chunk_size_mb', '64', 'Chunk size for multipart uploads'],
        ['api.listen_address', '0.0.0.0:8443', 'Management API bind address'],
        ['monitoring.prometheus_port', '9090', 'Prometheus metrics exporter port'],
    ]
    for r, row_data in enumerate(config_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')

    doc.add_heading('3.2 Cloud Provider Setup', level=2)
    doc.add_paragraph(
        'Each cloud provider requires a dedicated credentials block in the configuration file. '
        'CloudSync validates connectivity to all configured providers at startup.'
    )
    doc.add_paragraph(
        'For AWS, provide the access key ID and secret access key, or configure IAM role-based '
        'authentication when running on EC2 instances. For GCP, supply a service account JSON '
        'key file path. For Azure, use a storage account connection string or managed identity.'
    )

    # ===== PAGE 3: API Reference =====
    doc.add_heading('4. API Reference', level=1)
    doc.add_paragraph(
        'The CloudSync Management API provides RESTful endpoints for controlling sync operations, '
        'querying status, and managing configuration at runtime. All endpoints require TLS and '
        'bearer token authentication.'
    )

    doc.add_heading('4.1 Authentication', level=2)
    doc.add_paragraph(
        'Generate an API token using the CLI tool: cloudsync token generate --name admin-token '
        '--ttl 8760h. Include the token in the Authorization header of all requests.'
    )

    doc.add_heading('4.2 Sync Endpoints', level=2)

    endpoints_table = doc.add_table(rows=6, cols=3)
    endpoints_table.style = 'Table Grid'
    ep_headers = ['Method', 'Endpoint', 'Description']
    for i, h in enumerate(ep_headers):
        cell = endpoints_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    ep_data = [
        ['GET', '/api/v1/sync/status', 'Retrieve current sync status for all providers'],
        ['POST', '/api/v1/sync/trigger', 'Manually trigger a sync cycle'],
        ['GET', '/api/v1/sync/history', 'List completed sync operations with timestamps'],
        ['DELETE', '/api/v1/sync/queue/{id}', 'Cancel a pending sync job'],
        ['PUT', '/api/v1/config/reload', 'Hot-reload configuration without restart'],
    ]
    for r, row_data in enumerate(ep_data, 1):
        for c, val in enumerate(row_data):
            endpoints_table.cell(r, c).text = val

    doc.add_paragraph('')

    doc.add_heading('4.3 Monitoring Endpoints', level=2)
    doc.add_paragraph(
        'CloudSync exposes Prometheus-compatible metrics at /metrics on the configured monitoring '
        'port. Key metrics include cloudsync_sync_duration_seconds, cloudsync_files_transferred_total, '
        'cloudsync_errors_total, and cloudsync_queue_depth.'
    )
    doc.add_paragraph(
        'A Grafana dashboard template is included in the distribution under /opt/cloudsync/dashboards/. '
        'Import the JSON file into your Grafana instance and configure the Prometheus data source.'
    )

    doc.add_heading('4.4 Webhooks', level=2)
    doc.add_paragraph(
        'Configure webhook notifications for sync events by adding webhook entries to the '
        'configuration file. Each webhook supports filtering by event type (sync_complete, '
        'sync_failed, conflict_detected) and target URL with optional custom headers.'
    )

    # ===== PAGE 4: Monitoring & Security =====
    doc.add_heading('5. Monitoring and Alerting', level=1)
    doc.add_paragraph(
        'Effective monitoring is critical for maintaining reliable sync operations. CloudSync '
        'provides multiple monitoring channels including structured logging, Prometheus metrics, '
        'and webhook-based alerts.'
    )

    doc.add_heading('5.1 Log Management', level=2)
    doc.add_paragraph(
        'CloudSync writes structured JSON logs to /var/log/cloudsync/. Log rotation is handled '
        'by the built-in rotation mechanism with configurable retention. Logs include correlation '
        'IDs for tracing individual sync operations across components.'
    )
    doc.add_paragraph(
        'Log levels can be adjusted at runtime via the API without restarting the service. '
        'Available levels are DEBUG, INFO, WARN, and ERROR. Production environments should '
        'use INFO level with DEBUG enabled only for specific subsystems during investigation.'
    )

    doc.add_heading('5.2 Alert Rules', level=2)
    doc.add_paragraph(
        'The following Prometheus alert rules are recommended for production deployments:'
    )
    alerts = [
        'SyncFailureRate — fires when error rate exceeds 5% over 15 minutes',
        'SyncLatencyHigh — fires when p99 sync duration exceeds 300 seconds',
        'QueueBacklog — fires when queue depth exceeds 1000 pending jobs',
        'StorageQuotaWarning — fires when any provider reaches 90% storage utilization',
        'CertificateExpiringSoon — fires 30 days before TLS certificate expiration',
    ]
    for alert in alerts:
        doc.add_paragraph(alert, style='List Bullet')

    doc.add_heading('6. Security', level=1)
    doc.add_paragraph(
        'CloudSync implements defense-in-depth security controls including encryption at rest '
        'and in transit, role-based access control, audit logging, and network isolation. All '
        'inter-component communication uses mutual TLS (mTLS) with automatically rotated certificates.'
    )

    doc.add_heading('6.1 Access Control', level=2)
    doc.add_paragraph(
        'The platform supports three built-in roles: Admin (full access), Operator (sync management '
        'and monitoring), and Viewer (read-only status access). Custom roles can be defined using '
        'fine-grained permission policies in the access control configuration block.'
    )

    doc.add_heading('6.2 Audit Trail', level=2)
    doc.add_paragraph(
        'All API operations and configuration changes are recorded in an immutable audit log '
        'stored in the PostgreSQL metadata database. Audit entries include the authenticated '
        'principal, action performed, resource affected, and timestamp. The audit log supports '
        'export to external SIEM systems via Syslog or webhook integration.'
    )

    # ===== PAGE 5: Troubleshooting (target heading) =====
    doc.add_heading('7. Troubleshooting', level=1)
    doc.add_paragraph(
        'This section provides guidance for diagnosing and resolving common issues encountered '
        'during CloudSync operation. For issues not covered here, contact the support team with '
        'the correlation ID from the relevant log entries.'
    )

    doc.add_heading('7.1 Sync Failures', level=2)
    doc.add_paragraph(
        'If sync operations are failing consistently, check the following:'
    )
    sync_checks = [
        'Verify network connectivity to the target cloud provider endpoints using curl or openssl s_client',
        'Confirm that the credentials in config.yaml are valid and have not expired',
        'Check for storage quota exhaustion on the destination provider',
        'Review the sync error logs at /var/log/cloudsync/sync-errors.log for detailed failure reasons',
        'Ensure the PostgreSQL metadata database is accessible and has sufficient disk space',
    ]
    for check in sync_checks:
        doc.add_paragraph(check, style='List Bullet')

    doc.add_heading('7.2 Performance Degradation', level=2)
    doc.add_paragraph(
        'Performance issues typically stem from resource contention or suboptimal configuration. '
        'Start by reviewing the Prometheus dashboard for CPU, memory, and I/O utilization. If the '
        'sync queue depth is consistently growing, consider increasing sync.max_concurrent_jobs or '
        'deploying additional worker nodes.'
    )
    doc.add_paragraph(
        'For large file transfers exceeding 1 GB, verify that storage.chunk_size_mb is set to at '
        'least 64 MB. Smaller chunk sizes generate excessive API calls and reduce throughput. '
        'Network bandwidth can be verified using iperf3 between the CloudSync host and cloud endpoints.'
    )

    doc.add_heading('7.3 Database Issues', level=2)
    doc.add_paragraph(
        'CloudSync relies on PostgreSQL for metadata, job scheduling, and audit logs. Common '
        'database-related issues include connection pool exhaustion (increase db.max_connections), '
        'slow queries from large audit tables (run the provided maintenance SQL script monthly), '
        'and replication lag in HA setups (check pg_stat_replication).'
    )

    doc.add_heading('7.4 Certificate Errors', level=2)
    doc.add_paragraph(
        'TLS certificate errors typically manifest as "x509: certificate signed by unknown authority" '
        'in the logs. Ensure the CA certificate chain is complete in /etc/cloudsync/certs/ca.pem. '
        'For self-signed certificates in development environments, set tls.skip_verify to true '
        '(never in production).'
    )

    # ===== Appendix =====
    doc.add_heading('8. Appendix', level=1)

    doc.add_heading('8.1 Environment Variables', level=2)
    env_table = doc.add_table(rows=6, cols=2)
    env_table.style = 'Table Grid'
    env_table.cell(0, 0).text = 'Variable'
    env_table.cell(0, 1).text = 'Purpose'
    for run in env_table.cell(0, 0).paragraphs[0].runs:
        run.bold = True
    for run in env_table.cell(0, 1).paragraphs[0].runs:
        run.bold = True

    env_data = [
        ['CLOUDSYNC_CONFIG_PATH', 'Override default configuration file path'],
        ['CLOUDSYNC_LOG_LEVEL', 'Set logging verbosity (DEBUG, INFO, WARN, ERROR)'],
        ['CLOUDSYNC_DB_URL', 'PostgreSQL connection string override'],
        ['CLOUDSYNC_ENCRYPTION_KEY', 'Master encryption key (base64-encoded)'],
        ['CLOUDSYNC_LICENSE_KEY', 'Enterprise license activation key'],
    ]
    for r, row_data in enumerate(env_data, 1):
        for c, val in enumerate(row_data):
            env_table.cell(r, c).text = val

    doc.add_heading('8.2 Changelog', level=2)
    changelog = [
        'v3.2.1 (2026-03-28) — Fixed race condition in concurrent chunk uploads to Azure',
        'v3.2.0 (2026-02-15) — Added webhook notifications and custom alert rules',
        'v3.1.0 (2025-12-01) — Introduced GCP Dual-Region storage support',
        'v3.0.0 (2025-09-15) — Major release with new API v1, PostgreSQL 15 requirement',
    ]
    for entry in changelog:
        doc.add_paragraph(entry, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
