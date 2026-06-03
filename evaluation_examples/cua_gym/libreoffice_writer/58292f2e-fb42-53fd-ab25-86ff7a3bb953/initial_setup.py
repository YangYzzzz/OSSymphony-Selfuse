"""
Initial Setup: Footer margin too close to body text
Task ID: writer_fs_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    # Bottom margin small so footer is close to body text
    # Initial: spacing ~0.2cm, height ~0.5cm -> bottom_margin ~0.7cm
    section.bottom_margin = Cm(1.0)

    # --- Enable footer with content ---
    footer = section.footer
    footer.is_linked_to_previous = False

    # Footer distance from page edge (controls where footer sits)
    # Small value = footer close to page bottom, leaving little space
    section.footer_distance = Cm(0.5)

    # Add footer content
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fp.add_run("Quarterly Performance Report - Confidential")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Add page number to footer
    run_space = fp.add_run("  |  Page ")
    run_space.font.size = Pt(9)
    run_space.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    # Page number field
    r1 = fp.add_run()
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    r2.font.size = Pt(9)
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3.font.size = Pt(9)
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    # --- Document Content ---
    # Title
    title = doc.add_heading("Quarterly Performance Report - Q1 2025", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = subtitle.add_run("Prepared by the Strategic Planning Division")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph("")  # spacer

    # Section 1
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "The first quarter of 2025 has shown remarkable progress across all major "
        "business units. Total revenue reached $12.8 million, representing a 15.3% "
        "increase over the same period last year. The Engineering division led growth "
        "with a 22% improvement in productivity metrics, while the Marketing team "
        "successfully launched three new campaigns that exceeded their KPI targets."
    )
    doc.add_paragraph(
        "Operating expenses remained within budget at $9.2 million, yielding an "
        "operating margin of 28.1%. This represents the highest quarterly margin "
        "achieved since the company's restructuring in 2023. Cash reserves stand "
        "at $4.5 million, providing adequate runway for the planned expansion "
        "into the Asia-Pacific market scheduled for Q3."
    )

    # Section 2
    doc.add_heading("Department Highlights", level=2)

    doc.add_heading("Engineering", level=3)
    doc.add_paragraph(
        "Under the leadership of Director Sarah Chen, the Engineering team completed "
        "the migration to the new cloud infrastructure ahead of schedule. Key "
        "achievements include reducing API response times by 40%, deploying the "
        "automated testing pipeline, and onboarding 12 new engineers across the "
        "San Francisco and Berlin offices."
    )

    doc.add_heading("Marketing & Sales", level=3)
    doc.add_paragraph(
        "Marcus Johnson's team delivered an outstanding quarter with the successful "
        "launch of the 'Innovate Forward' brand campaign. Social media engagement "
        "increased by 67%, and the sales pipeline grew to $8.3 million in qualified "
        "opportunities. The team also secured partnerships with three Fortune 500 "
        "companies for co-marketing initiatives."
    )

    doc.add_heading("Human Resources", level=3)
    doc.add_paragraph(
        "The HR department, led by Priya Patel, implemented the new employee "
        "wellness program that has already seen 78% participation. Retention rates "
        "improved to 94.2%, up from 89.5% in Q4 2024. The annual compensation "
        "review was completed on time, with an average salary adjustment of 4.2% "
        "across all departments."
    )

    # Section 3 - Table
    doc.add_heading("Financial Summary", level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"

    headers = ["Metric", "Q1 2025", "Q4 2024", "Change"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ["Total Revenue", "$12.8M", "$11.4M", "+12.3%"],
        ["Operating Expenses", "$9.2M", "$9.0M", "+2.2%"],
        ["Net Income", "$3.6M", "$2.4M", "+50.0%"],
        ["Employee Count", "342", "318", "+7.5%"],
        ["Customer Satisfaction", "4.7/5.0", "4.5/5.0", "+4.4%"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # Section 4
    doc.add_paragraph("")
    doc.add_heading("Strategic Outlook", level=2)
    doc.add_paragraph(
        "Looking ahead to Q2 2025, the company is well-positioned for continued "
        "growth. Priority initiatives include the launch of Version 3.0 of our "
        "flagship product, expansion of the customer success team, and the "
        "establishment of a regional office in Singapore. The board has approved "
        "a capital expenditure budget of $2.1 million for these initiatives."
    )
    doc.add_paragraph(
        "Risk factors to monitor include potential supply chain disruptions in "
        "the semiconductor sector, evolving regulatory requirements in the EU "
        "market, and competitive pressure from two new entrants in our primary "
        "market segment. Mitigation strategies have been developed for each "
        "identified risk and will be reviewed at the next board meeting on "
        "April 18, 2025."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
