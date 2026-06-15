"""
Initial Setup: Create an event flyer document without any footer banner shape.
Task ID: writer_frd_073
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # --- Title ---
    title = doc.add_heading('Annual Community Arts Festival 2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8F)
        run.font.size = Pt(28)

    # --- Subtitle / Tagline ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Celebrating Creativity in Our Neighborhood')
    run.font.size = Pt(16)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Blank spacer ---
    doc.add_paragraph()

    # --- Event Details ---
    details_heading = doc.add_heading('Event Details', level=1)
    for run in details_heading.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    details = [
        ('Date:', 'Saturday, June 14, 2026'),
        ('Time:', '10:00 AM - 6:00 PM'),
        ('Venue:', 'Riverside Park Amphitheater, 450 Lakewood Blvd'),
        ('Admission:', 'Free for all ages'),
    ]
    for label, value in details:
        para = doc.add_paragraph()
        run_label = para.add_run(label + ' ')
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_value = para.add_run(value)
        run_value.font.size = Pt(12)

    doc.add_paragraph()

    # --- What to Expect ---
    expect_heading = doc.add_heading('What to Expect', level=1)
    for run in expect_heading.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    bullets = [
        'Live performances by local musicians and dance groups',
        'Interactive art workshops for children and adults',
        'Artisan market featuring handmade crafts and jewelry',
        'Food trucks and local restaurant pop-ups',
        'Outdoor sculpture exhibit curated by the Riverside Gallery',
        'Face painting, balloon artists, and family activities',
    ]
    for item in bullets:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()

    # --- Featured Artists ---
    artists_heading = doc.add_heading('Featured Artists', level=1)
    for run in artists_heading.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    artists = [
        ('Elena Vasquez', 'Abstract Oil Paintings - Winner of the 2025 Regional Arts Prize'),
        ('Marcus Okafor', 'Contemporary Sculpture Installation'),
        ('Priya Mehta', 'Traditional Henna Art & Live Demonstration'),
        ('James Whitfield', 'Acoustic Folk Performance (Main Stage, 2:00 PM)'),
    ]
    for name, desc in artists:
        para = doc.add_paragraph()
        run_name = para.add_run(name + ' - ')
        run_name.bold = True
        run_name.font.size = Pt(11)
        run_desc = para.add_run(desc)
        run_desc.font.size = Pt(11)

    doc.add_paragraph()

    # --- Sponsors ---
    sponsors_heading = doc.add_heading('Our Sponsors', level=2)
    for run in sponsors_heading.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    sponsors_para = doc.add_paragraph()
    sponsors_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sponsors_para.add_run(
        'Greenfield Community Foundation  |  Lakewood Chamber of Commerce  |  '
        'Riverside Gallery  |  BlueStar Financial Group'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # --- Contact Info ---
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = contact.add_run('For more information, visit www.riversideartsfest.org or call (555) 234-7890')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
