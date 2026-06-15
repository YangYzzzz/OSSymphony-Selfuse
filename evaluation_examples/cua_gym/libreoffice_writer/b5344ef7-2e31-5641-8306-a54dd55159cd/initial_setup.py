"""
Initial Setup: Create a business letter document (no envelope configured)
Task ID: writer_biz_056
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_056'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup: standard US Letter ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Sender letterhead ---
    sender_block = doc.add_paragraph()
    sender_block.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sender_block.paragraph_format.space_after = Pt(0)
    run = sender_block.add_run("Meridian Solutions Inc.")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    addr_lines = [
        "1200 Commerce Drive, Suite 400",
        "Chicago, IL 60601",
        "Phone: (312) 555-8200",
        "Email: info@meridiansolutions.com",
    ]
    for line in addr_lines:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(24)
    date_para.paragraph_format.space_after = Pt(12)
    r = date_para.add_run("March 28, 2026")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # --- Recipient address ---
    recipient_lines = [
        "Mr. James Wilson",
        "Pinnacle Corp",
        "500 Park Avenue",
        "New York, NY 10022",
    ]
    for line in recipient_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # --- Salutation ---
    sal = doc.add_paragraph()
    sal.paragraph_format.space_before = Pt(12)
    sal.paragraph_format.space_after = Pt(6)
    r = sal.add_run("Dear Mr. Wilson,")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # --- Body paragraphs ---
    body_texts = [
        "Thank you for meeting with our team last Thursday to discuss the upcoming "
        "digital transformation initiative for Pinnacle Corp. We were impressed by your "
        "organization\u2019s commitment to modernizing its technology infrastructure and "
        "believe Meridian Solutions is well-positioned to support this effort.",

        "As discussed, our proposal covers three key areas: migration of your legacy ERP "
        "system to a cloud-based platform, implementation of an integrated data analytics "
        "dashboard for real-time business intelligence, and a comprehensive cybersecurity "
        "audit of your existing network architecture. The total estimated investment for "
        "Phase 1 is $285,000, with a projected timeline of 14 weeks from contract signing.",

        "Our project manager, Elena Vasquez, will serve as your primary point of contact "
        "throughout the engagement. She has over twelve years of experience leading enterprise "
        "transformation projects and has worked extensively with organizations in the financial "
        "services sector.",

        "We have attached the detailed Statement of Work (SOW) and the Master Services "
        "Agreement (MSA) for your legal team\u2019s review. Please do not hesitate to reach "
        "out if you have any questions or require any modifications to the proposed scope.",
    ]
    for text in body_texts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(0)
    r = closing.add_run("Sincerely,")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    # --- Signature block ---
    sig_lines = [
        "",
        "",
        "David Nakamura",
        "Senior Vice President, Client Services",
        "Meridian Solutions Inc.",
    ]
    for line in sig_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        if line == "David Nakamura":
            r.bold = True

    # --- Enclosure note ---
    enc = doc.add_paragraph()
    enc.paragraph_format.space_before = Pt(12)
    r = enc.add_run("Enclosures: Statement of Work (SOW), Master Services Agreement (MSA)")
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    r.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
