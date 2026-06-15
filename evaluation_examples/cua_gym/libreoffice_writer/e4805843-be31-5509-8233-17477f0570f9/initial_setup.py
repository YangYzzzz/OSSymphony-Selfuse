"""
Initial Setup: Convert tab-separated text block into a proper table
Task ID: writer_tm_010
Domain: libreoffice_writer

Creates a document with tab-separated text data (no table).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title heading
    doc.add_heading("Raw Data", level=1)

    # Add an introductory paragraph
    doc.add_paragraph(
        "The following data was exported from our regional office database. "
        "Please format it appropriately for the quarterly report."
    )

    # Add the tab-separated text block as individual paragraphs
    # Each line is a paragraph with tab-separated values
    lines = [
        "Name\tAge\tCity",
        "Alice\t30\tBoston",
        "Bob\t25\tDenver",
        "Carol\t35\tSeattle",
    ]
    for line in lines:
        para = doc.add_paragraph(line)
        # Use a monospace-ish font to make it look like raw data
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(11)

    # Add a closing paragraph
    doc.add_paragraph(
        "End of data export. Please convert the above into a structured format."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
