"""
Reward Script: Convert tab-separated text block into a proper table
Task ID: writer_tm_010
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table exists with correct dimensions (4 rows x 3 cols)
  Component 2 (0.3): Header row has correct values (Name, Age, City)
  Component 3 (0.3): Data rows have correct values
  Component 4 (0.1): Tab-separated text paragraphs removed from body
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_010'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions — 4 rows x 3 cols (0.3 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 4 and num_cols == 3:
                print(f"PASS: Component 1 — Table found with {num_rows} rows x {num_cols} cols (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 1 — Table has {num_rows} rows x {num_cols} cols, expected 4x3")
        else:
            print(f"FAIL: Component 1 — No tables found in document (found {len(doc.tables)})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row contains correct values (0.3 points)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            header_cells = [table.cell(0, c).text.strip() for c in range(min(len(table.columns), 3))]
            expected_header = ['Name', 'Age', 'City']
            if header_cells == expected_header:
                print(f"PASS: Component 2 — Header row correct: {header_cells} (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Header row: {header_cells}, expected {expected_header}")
        else:
            print("FAIL: Component 2 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Data rows contain correct values (0.3 points)
    # Each correct row earns 0.1 points (3 data rows x 0.1 = 0.3)
    expected_data = [
        ['Alice', '30', 'Boston'],
        ['Bob', '25', 'Denver'],
        ['Carol', '35', 'Seattle'],
    ]
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            data_score = 0.0
            for row_idx, expected_row in enumerate(expected_data):
                actual_row_idx = row_idx + 1  # skip header
                if actual_row_idx < len(table.rows):
                    actual_cells = [table.cell(actual_row_idx, c).text.strip() for c in range(min(len(table.columns), 3))]
                    if actual_cells == expected_row:
                        print(f"PASS: Component 3.{row_idx+1} — Row {actual_row_idx} correct: {actual_cells} (0.1 pts)")
                        data_score += 0.1
                    else:
                        print(f"FAIL: Component 3.{row_idx+1} — Row {actual_row_idx}: {actual_cells}, expected {expected_row}")
                else:
                    print(f"FAIL: Component 3.{row_idx+1} — Row {actual_row_idx} does not exist")
            if data_score > 0:
                total_score += data_score
        else:
            print("FAIL: Component 3 — No tables found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Tab-separated text paragraphs removed (0.1 points)
    # In the initial file, paragraphs like "Name\tAge\tCity" exist as plain text.
    # After conversion, these tab-separated lines should no longer be plain paragraphs.
    try:
        tab_lines_found = []
        tab_data_patterns = ['Name\tAge\tCity', 'Alice\t30\tBoston', 'Bob\t25\tDenver', 'Carol\t35\tSeattle']
        for para in doc.paragraphs:
            for pattern in tab_data_patterns:
                if pattern in para.text:
                    tab_lines_found.append(para.text)
        if len(tab_lines_found) == 0:
            print(f"PASS: Component 4 — No tab-separated data paragraphs remain (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 4 — {len(tab_lines_found)} tab-separated paragraph(s) still present: {tab_lines_found}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 1)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
