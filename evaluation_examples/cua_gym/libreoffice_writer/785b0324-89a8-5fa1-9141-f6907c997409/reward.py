"""
Reward Script: Set header spacing to 0.5cm and add 'Q3 Financial Review' centered in header.
Task ID: writer_page_025
Domain: libreoffice_writer
Scoring:
  Component 1: Header paragraph text is 'Q3 Financial Review' (0.5 pts)
  Component 2: Header paragraph alignment is CENTER (0.25 pts)
  Component 3: Header distance (from top of page) is approximately 0.5cm (within 0.1cm tolerance) (0.25 pts)
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_025'
FILE_NAME = 'financial_review.docx'

# 0.5cm expressed in EMU (English Metric Units)
# 1 inch = 914400 EMU, 1 cm = 914400/2.54 EMU
CM_TO_EMU = 914400 / 2.54
TARGET_HEADER_DISTANCE_CM = 0.5
TOLERANCE_CM = 0.1  # accept values within ±0.1cm of 0.5cm


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document (precondition gate — if corrupt, return 0.0 immediately)
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract the first section's header (all changes are in section 0)
    try:
        section = doc.sections[0]
        header = section.header
    except Exception as e:
        print(f"CRITICAL: Cannot access section/header: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Header paragraph text is 'Q3 Financial Review' (0.5 pts)
    # FAILS on initial (header is empty), PASSES on golden (text present)
    try:
        header_text = ""
        if header.paragraphs:
            # Collect all text from header paragraphs
            header_text = "\n".join(p.text for p in header.paragraphs).strip()

        expected_text = "Q3 Financial Review"
        if header_text == expected_text:
            print(f"PASS: Component 1 — Header text is '{header_text}' (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Expected header text '{expected_text}', found '{header_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — Could not check header text: {e}")

    # Component 2: Header paragraph alignment is CENTER (0.25 pts)
    # FAILS on initial (alignment is None / not CENTER), PASSES on golden (CENTER)
    try:
        header_para = header.paragraphs[0] if header.paragraphs else None
        if header_para is not None:
            align = header_para.paragraph_format.alignment
            if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                print(f"PASS: Component 2 — Header paragraph alignment is CENTER (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — Expected CENTER alignment, found {align}")
        else:
            print(f"FAIL: Component 2 — No header paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 2 — Could not check header alignment: {e}")

    # Component 3: Header distance is approximately 0.5cm (±0.1cm tolerance) (0.25 pts)
    # FAILS on initial (~1.25cm), PASSES on golden (~0.5cm)
    try:
        header_distance_emu = section.header_distance
        header_distance_cm = header_distance_emu / CM_TO_EMU
        diff_cm = abs(header_distance_cm - TARGET_HEADER_DISTANCE_CM)
        if diff_cm <= TOLERANCE_CM:
            print(f"PASS: Component 3 — Header distance is {header_distance_cm:.4f}cm "
                  f"(target {TARGET_HEADER_DISTANCE_CM}cm, diff={diff_cm:.4f}cm) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Header distance is {header_distance_cm:.4f}cm, "
                  f"expected ~{TARGET_HEADER_DISTANCE_CM}cm (tolerance ±{TOLERANCE_CM}cm), "
                  f"diff={diff_cm:.4f}cm")
    except Exception as e:
        print(f"ERROR: Component 3 — Could not check header distance: {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path on the VM
file_path = os.path.join(WORKDIR, FILE_NAME)
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
