"""
Initial Setup: Create a 5-page quarterly financial review document with empty header
Task ID: writer_page_025
Domain: libreoffice_writer

Creates financial_review.docx on ~/Desktop/ with:
- A4 portrait, margins: top=2.54cm, bottom=2.54cm, left=2.54cm, right=2.54cm
- Header ENABLED but EMPTY (no header text)
- Default header distance (~1.25cm)
- 5 pages of realistic quarterly financial review content
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'financial_review'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_para(doc, text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    para = doc.add_paragraph(text)
    para.paragraph_format.alignment = alignment
    return para


def add_table_row(table, row_idx, values):
    row = table.rows[row_idx]
    for col_idx, val in enumerate(values):
        row.cells[col_idx].text = str(val)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Configure page settings: A4 portrait
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Enable header but keep it EMPTY (no text)
    # Default header_distance (~1.25cm) — do NOT set to 0.5cm (that's the task)
    section.header_distance = Cm(1.25)
    header = section.header
    header.is_linked_to_previous = False
    # Ensure header paragraph exists but is empty
    if header.paragraphs:
        header.paragraphs[0].clear()
    else:
        header.add_paragraph()

    # ---- PAGE 1: Executive Summary ----
    add_heading(doc, 'Q3 2025 Financial Review', level=0)
    add_para(doc, 'Prepared by: Finance Department | Meridian Global Partners LLC')
    add_para(doc, 'Report Date: October 15, 2025')
    doc.add_paragraph()

    add_heading(doc, 'Executive Summary', level=1)
    add_para(doc,
        'Meridian Global Partners LLC delivered a strong third quarter performance, '
        'driven by robust revenue growth in our core technology services segment and '
        'disciplined cost management across all divisions. Total consolidated revenue '
        'for Q3 2025 reached $142.8 million, representing a 12.4% increase year-over-year '
        'and exceeding the initial guidance of $138 million.'
    )
    add_para(doc,
        'Operating income improved to $28.6 million (20.0% margin), up from $24.1 million '
        'in Q3 2024. Net income attributable to common shareholders was $21.3 million, '
        'or $1.47 per diluted share, compared to $17.9 million, or $1.24 per diluted share, '
        'in the prior year period.'
    )
    add_para(doc,
        'Free cash flow generation remained strong at $19.2 million during the quarter, '
        'supporting our ongoing share repurchase program and strategic investment initiatives. '
        'The Board of Directors approved a quarterly dividend of $0.22 per share, payable '
        'November 15, 2025, to shareholders of record as of November 1, 2025.'
    )
    doc.add_page_break()

    # ---- PAGE 2: Revenue Analysis ----
    add_heading(doc, 'Revenue Analysis', level=1)
    add_para(doc,
        'The following table summarizes revenue by business segment for Q3 2025 compared '
        'to Q3 2024 and the preceding quarter Q2 2025.'
    )

    # Revenue table
    rev_table = doc.add_table(rows=6, cols=4)
    rev_table.style = 'Table Grid'
    headers = ['Business Segment', 'Q3 2025 ($M)', 'Q3 2024 ($M)', 'YoY Change']
    for col_idx, h in enumerate(headers):
        cell = rev_table.rows[0].cells[col_idx]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    data_rows = [
        ['Technology Services', '58.4', '51.2', '+14.1%'],
        ['Enterprise Solutions', '41.7', '37.8', '+10.3%'],
        ['Cloud Infrastructure', '29.3', '23.6', '+24.2%'],
        ['Professional Services', '13.4', '14.1', '-5.0%'],
        ['Total Revenue', '142.8', '126.7', '+12.4%'],
    ]
    for row_idx, row_data in enumerate(data_rows, 1):
        for col_idx, val in enumerate(row_data):
            rev_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    add_para(doc,
        'Technology Services continued to be our largest revenue contributor, accounting for '
        '40.9% of total revenue. The segment benefited from the expansion of managed services '
        'contracts with enterprise clients in the financial services and healthcare sectors. '
        'Cloud Infrastructure posted the strongest growth at 24.2% YoY, reflecting increased '
        'demand for hybrid cloud deployment solutions.'
    )
    add_para(doc,
        'Professional Services revenue declined 5.0% YoY, primarily due to the conclusion of '
        'two large consulting engagements in Q2 2025. The pipeline for new consulting projects '
        'remains healthy, with $18.5 million in signed but not yet commenced engagements as of '
        'September 30, 2025.'
    )
    doc.add_page_break()

    # ---- PAGE 3: Operating Expenses ----
    add_heading(doc, 'Operating Expenses and Margin Analysis', level=1)
    add_para(doc,
        'Total operating expenses for Q3 2025 were $114.2 million, compared to $102.6 million '
        'in Q3 2024, representing an 11.3% increase — lower than revenue growth of 12.4%, '
        'demonstrating continued operating leverage.'
    )

    # Expense breakdown table
    exp_table = doc.add_table(rows=6, cols=3)
    exp_table.style = 'Table Grid'
    exp_headers = ['Expense Category', 'Q3 2025 ($M)', 'Q3 2024 ($M)']
    for col_idx, h in enumerate(exp_headers):
        cell = exp_table.rows[0].cells[col_idx]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    exp_data = [
        ['Cost of Revenue', '67.3', '60.8'],
        ['Research & Development', '18.9', '17.2'],
        ['Sales & Marketing', '15.6', '13.9'],
        ['General & Administrative', '12.4', '10.7'],
        ['Total Operating Expenses', '114.2', '102.6'],
    ]
    for row_idx, row_data in enumerate(exp_data, 1):
        for col_idx, val in enumerate(row_data):
            exp_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    add_para(doc,
        'Gross margin improved to 52.8% in Q3 2025, up from 52.0% in Q3 2024. The improvement '
        'reflects favorable revenue mix shifts toward higher-margin Cloud Infrastructure services '
        'and efficiency gains in our delivery operations. R&D investment increased 9.9% YoY '
        'as we accelerated development of our next-generation analytics platform, scheduled '
        'for general availability in Q1 2026.'
    )
    add_para(doc,
        'Sales and marketing expenses grew 12.2% YoY, in line with revenue growth, as we '
        'expanded our go-to-market teams in the Asia-Pacific region. G&A expenses increased '
        '15.9% YoY, primarily due to costs associated with our enterprise resource planning '
        'system upgrade completed in August 2025.'
    )
    doc.add_page_break()

    # ---- PAGE 4: Balance Sheet & Cash Flow ----
    add_heading(doc, 'Balance Sheet Highlights', level=1)
    add_para(doc,
        'As of September 30, 2025, total assets were $483.6 million, compared to $447.2 million '
        'at June 30, 2025. The increase primarily reflects higher accounts receivable balances '
        'corresponding to strong quarter-end revenue, partially offset by reduction in cash '
        'balances following the $15.0 million debt repayment in September.'
    )

    bs_table = doc.add_table(rows=7, cols=3)
    bs_table.style = 'Table Grid'
    bs_headers = ['Balance Sheet Item', 'Sep 30, 2025 ($M)', 'Jun 30, 2025 ($M)']
    for col_idx, h in enumerate(bs_headers):
        cell = bs_table.rows[0].cells[col_idx]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    bs_data = [
        ['Cash & Equivalents', '87.4', '98.3'],
        ['Accounts Receivable', '124.7', '108.2'],
        ['Total Current Assets', '243.8', '231.5'],
        ['Total Assets', '483.6', '447.2'],
        ['Total Liabilities', '198.4', '209.7'],
        ["Total Shareholders' Equity", '285.2', '237.5'],
    ]
    for row_idx, row_data in enumerate(bs_data, 1):
        for col_idx, val in enumerate(row_data):
            bs_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    add_heading(doc, 'Cash Flow Summary', level=2)
    add_para(doc,
        'Operating cash flow for Q3 2025 was $22.8 million. Capital expenditures totaled '
        '$3.6 million, primarily for data center equipment upgrades. Free cash flow of '
        '$19.2 million funded $8.4 million in share repurchases and the $15.0 million '
        'scheduled debt repayment, with the remainder increasing the revolver balance.'
    )
    doc.add_page_break()

    # ---- PAGE 5: Outlook & Guidance ----
    add_heading(doc, 'Q4 2025 Outlook and Full-Year Guidance', level=1)
    add_para(doc,
        'Management reaffirms full-year 2025 revenue guidance in the range of $555 million '
        'to $562 million, representing 11% to 12% growth compared to fiscal 2024. This guidance '
        'reflects anticipated sequential improvement in Professional Services revenue in Q4, '
        'continued strong demand in Cloud Infrastructure, and typical Q4 seasonality effects '
        'in Enterprise Solutions.'
    )
    add_para(doc,
        'For Q4 2025 specifically, management expects:'
    )

    # Guidance bullet list
    guidance_items = [
        'Revenue: $147 million to $151 million',
        'Operating margin: 19.5% to 21.0%',
        'Diluted EPS: $1.51 to $1.58',
        'Free cash flow: $18 million to $22 million',
        'R&D investment: approximately $20 million',
    ]
    for item in guidance_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_heading(doc, 'Key Risks and Uncertainties', level=2)
    add_para(doc,
        'Forward-looking statements in this report are subject to significant risks and '
        'uncertainties, including: macroeconomic conditions affecting enterprise IT spending; '
        'competitive pricing pressures in the Cloud Infrastructure market; foreign exchange '
        'rate fluctuations impacting APAC revenue; and potential delays in the Q1 2026 '
        'analytics platform launch.'
    )
    add_para(doc,
        'This report contains forward-looking statements within the meaning of the Private '
        'Securities Litigation Reform Act of 1995. Actual results may differ materially from '
        'those projected. Please refer to the Risk Factors section of our most recent Annual '
        'Report on Form 10-K for a description of factors that could cause actual results '
        'to differ.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
