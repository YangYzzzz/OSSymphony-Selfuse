"""
Reward Script: CS301 Database Systems Course Syllabus
Task ID: writer_wf_049
Domain: libreoffice_writer
Scoring:
  C1: University name "Pacific State University" in document (0.10)
  C2: Course title "CS301: Database Systems" present (0.10)
  C3: Course info table with 5 required fields (0.20)
  C4: Course Description section with paragraph text (0.10)
  C5: 5 numbered Learning Objectives (0.15)
  C6: Grading Policy table with 5 data rows summing to 100% (0.15)
  C7: Weekly Schedule table with 8 weeks and 4 columns (0.15)
  C8: Academic Integrity notice present (0.05)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_049'


def persist_app_state(domain: str):
    """Try to save any unsaved state in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_para_texts = [p.text.strip() for p in doc.paragraphs]
    all_text_lower = ' '.join(all_para_texts).lower()

    # Component 1: University name "Pacific State University" present (0.10 pts)
    try:
        found_university = False
        for p in doc.paragraphs:
            if 'pacific state university' in p.text.lower():
                found_university = True
                break
        # Also check section headers
        if not found_university:
            for s in doc.sections:
                if s.header and s.header.paragraphs:
                    for hp in s.header.paragraphs:
                        if 'pacific state university' in hp.text.lower():
                            found_university = True
                            break
        if found_university:
            print(f"PASS: Component 1 — 'Pacific State University' found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — 'Pacific State University' not found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Course title "CS301: Database Systems" present (0.10 pts)
    try:
        found_course = False
        for p in doc.paragraphs:
            if 'cs301' in p.text.lower() and 'database' in p.text.lower():
                found_course = True
                break
        if found_course:
            print(f"PASS: Component 2 — 'CS301: Database Systems' found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — 'CS301: Database Systems' not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Course info table with 5 required fields (0.20 pts)
    try:
        required_fields = ['course', 'instructor', 'office hours', 'email', 'prerequisites']
        found_info_table = False
        for table in doc.tables:
            # Check if this table has the required field labels
            field_texts = []
            for row in table.rows:
                first_cell = row.cells[0].text.strip().lower()
                field_texts.append(first_cell)
            matches = sum(1 for rf in required_fields if any(rf in ft for ft in field_texts))
            if matches >= 4:
                found_info_table = True
                # Check it has at least 5 rows
                if len(table.rows) >= 5 and len(table.columns) >= 2:
                    # Check that value cells are not empty
                    non_empty_values = sum(1 for row in table.rows if row.cells[1].text.strip())
                    if non_empty_values >= 4:
                        print(f"PASS: Component 3 — Course info table found with {len(table.rows)} rows, {matches}/5 fields, {non_empty_values} filled values (0.20 pts)")
                        total_score += 0.20
                    else:
                        print(f"FAIL: Component 3 — Course info table found but only {non_empty_values} values filled")
                else:
                    print(f"FAIL: Component 3 — Course info table found but dimensions wrong: {len(table.rows)}x{len(table.columns)}")
                break
        if not found_info_table:
            print(f"FAIL: Component 3 — No course info table found with required fields")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Course Description section with text (0.10 pts)
    try:
        found_desc = False
        for i, p in enumerate(doc.paragraphs):
            if 'course description' in p.text.lower():
                # Check there's a non-empty paragraph following
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    if doc.paragraphs[j].text.strip() and len(doc.paragraphs[j].text.strip()) > 20:
                        found_desc = True
                        break
                break
        if found_desc:
            print(f"PASS: Component 4 — Course Description section with content found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — Course Description section not found or empty")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 5 numbered Learning Objectives (0.15 pts)
    try:
        found_objectives_section = False
        objective_count = 0
        in_objectives = False
        for i, p in enumerate(doc.paragraphs):
            if 'learning objective' in p.text.lower():
                found_objectives_section = True
                in_objectives = True
                continue
            if in_objectives:
                # Check if this is still in objectives section (stop at next heading)
                if p.style and 'heading' in p.style.name.lower() and p.text.strip():
                    break
                # Count numbered items (List Number style or starts with digit)
                if p.text.strip():
                    if (p.style and 'list number' in p.style.name.lower()) or \
                       re.match(r'^\d+[\.\)]\s', p.text.strip()):
                        objective_count += 1
                    elif len(p.text.strip()) > 20:
                        # Could be a numbered item without explicit style
                        objective_count += 1

        if found_objectives_section and objective_count >= 5:
            print(f"PASS: Component 5 — {objective_count} learning objectives found (0.15 pts)")
            total_score += 0.15
        elif found_objectives_section and objective_count >= 3:
            partial = 0.15 * (objective_count / 5.0)
            print(f"PARTIAL: Component 5 — Only {objective_count}/5 objectives found ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Learning Objectives: section={'found' if found_objectives_section else 'missing'}, count={objective_count}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Grading Policy table with 5 data rows summing to 100% (0.15 pts)
    try:
        found_grading = False
        for table in doc.tables:
            # Look for a table with "Component" and "Weight" headers
            if len(table.rows) < 2:
                continue
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'component' in headers and 'weight' in headers:
                found_grading = True
                data_rows = len(table.rows) - 1  # minus header
                # Parse weights
                weight_col = headers.index('weight')
                total_weight = 0
                for ri in range(1, len(table.rows)):
                    w_text = table.rows[ri].cells[weight_col].text.strip()
                    m = re.search(r'(\d+)', w_text)
                    if m:
                        total_weight += int(m.group(1))

                if data_rows >= 5 and total_weight == 100:
                    print(f"PASS: Component 6 — Grading table has {data_rows} rows, weights sum to {total_weight}% (0.15 pts)")
                    total_score += 0.15
                elif data_rows >= 5 and 90 <= total_weight <= 110:
                    print(f"PARTIAL: Component 6 — Grading table has {data_rows} rows, weights sum to {total_weight}% (0.10 pts)")
                    total_score += 0.10
                else:
                    print(f"FAIL: Component 6 — Grading table: {data_rows} data rows, weights sum to {total_weight}%")
                break
        if not found_grading:
            print(f"FAIL: Component 6 — No grading policy table found with Component/Weight headers")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Weekly Schedule table with 8 weeks, 4 columns (0.15 pts)
    try:
        found_schedule = False
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'week' in headers and ('topic' in headers or 'subject' in headers):
                found_schedule = True
                data_rows = len(table.rows) - 1
                num_cols = len(table.columns)

                if data_rows >= 8 and num_cols >= 4:
                    # Check that week numbers are present
                    weeks_found = 0
                    for ri in range(1, len(table.rows)):
                        w_text = table.rows[ri].cells[0].text.strip()
                        if w_text and re.search(r'\d+', w_text):
                            weeks_found += 1
                    if weeks_found >= 8:
                        print(f"PASS: Component 7 — Schedule table: {data_rows} data rows, {num_cols} cols, {weeks_found} weeks (0.15 pts)")
                        total_score += 0.15
                    else:
                        partial = 0.15 * (weeks_found / 8.0)
                        print(f"PARTIAL: Component 7 — Schedule table: only {weeks_found}/8 weeks with numbers ({partial:.2f} pts)")
                        total_score += partial
                elif data_rows >= 6:
                    partial = 0.15 * (data_rows / 8.0)
                    print(f"PARTIAL: Component 7 — Schedule table: {data_rows}/8 rows, {num_cols} cols ({partial:.2f} pts)")
                    total_score += partial
                else:
                    print(f"FAIL: Component 7 — Schedule table too small: {data_rows} data rows, {num_cols} cols")
                break
        if not found_schedule:
            print(f"FAIL: Component 7 — No weekly schedule table found with Week/Topic headers")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: Academic Integrity notice present (0.05 pts)
    try:
        found_integrity = False
        for p in doc.paragraphs:
            if 'academic integrity' in p.text.lower():
                found_integrity = True
                break
        if found_integrity:
            print(f"PASS: Component 8 — Academic Integrity notice found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 8 — Academic Integrity notice not found")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    # Also check there is substantive content after "Academic Integrity" heading
    try:
        for i, p in enumerate(doc.paragraphs):
            if 'academic integrity' in p.text.lower() and p.style and 'heading' in p.style.name.lower():
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    if doc.paragraphs[j].text.strip() and len(doc.paragraphs[j].text.strip()) > 20:
                        print(f"  INFO: Academic Integrity has content paragraph")
                        break
                break
    except Exception:
        pass

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
