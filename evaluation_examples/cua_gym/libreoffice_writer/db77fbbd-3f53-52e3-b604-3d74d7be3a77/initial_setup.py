"""
Initial Setup: Create Annual Report document with Heading 2 sections
Task ID: writer_frd_012
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_012'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Meridian Technologies Annual Report 2025', level=1)

    # Introduction body
    doc.add_paragraph(
        'This annual report provides a comprehensive overview of Meridian Technologies\' '
        'performance, strategic initiatives, and financial results for the fiscal year ending '
        'December 31, 2025. Our commitment to innovation and sustainable growth continues to '
        'drive value for shareholders, employees, and the communities we serve.'
    )

    # Heading 2 sections with body text (8 Heading 2 paragraphs total)

    # 1
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'Fiscal year 2025 marked a transformative period for Meridian Technologies. '
        'Revenue grew 18.3% year-over-year to $4.72 billion, driven by strong demand in our '
        'cloud infrastructure and cybersecurity divisions. Operating margin expanded to 22.1%, '
        'reflecting disciplined cost management and favorable product mix shifts toward '
        'higher-margin recurring revenue streams.'
    )

    # 2
    doc.add_heading('Financial Performance', level=2)
    doc.add_paragraph(
        'Total revenue reached $4.72 billion, up from $3.99 billion in the prior year. '
        'Gross profit increased to $2.98 billion with a gross margin of 63.1%. Research and '
        'development expenditures totaled $612 million, representing 13.0% of revenue. Net '
        'income attributable to common shareholders was $847 million, or $6.32 per diluted share, '
        'compared to $691 million, or $5.18 per diluted share, in the prior year.'
    )

    # 3
    doc.add_heading('Product Innovation and Development', level=2)
    doc.add_paragraph(
        'Our engineering teams delivered 14 major product releases across three divisions. '
        'The launch of CloudShield Pro in Q2 generated $238 million in first-year bookings, '
        'exceeding internal projections by 42%. The Meridian Edge computing platform achieved '
        'general availability in September, securing 73 enterprise customers within its first '
        'quarter. Patent filings increased 31% to 287 new applications.'
    )

    # 4
    doc.add_heading('Market Expansion and Customer Growth', level=2)
    doc.add_paragraph(
        'We expanded operations into six new markets across Southeast Asia and Latin America. '
        'Our total customer base grew to 12,400 enterprise accounts, a net increase of 2,150 '
        'from the prior year. Customer retention rate remained strong at 94.7%, and annual '
        'recurring revenue from existing customers grew 21.6% through upsell and cross-sell '
        'initiatives.'
    )

    # 5
    doc.add_heading('Sustainability and Corporate Responsibility', level=2)
    doc.add_paragraph(
        'Meridian Technologies achieved carbon neutrality across all Scope 1 and Scope 2 '
        'emissions for the second consecutive year. We invested $45 million in renewable energy '
        'infrastructure, bringing our data center renewable energy usage to 87%. The Meridian '
        'Foundation distributed $18.2 million in grants supporting STEM education programs '
        'in underserved communities across 14 countries.'
    )

    # 6
    doc.add_heading('Talent and Organizational Development', level=2)
    doc.add_paragraph(
        'Our global workforce expanded to 28,700 employees, with 4,300 new hires across '
        'engineering, sales, and customer success functions. Employee engagement scores reached '
        '82%, up from 78% in the prior year. We launched the Meridian Leadership Academy, '
        'enrolling 640 high-potential managers in an intensive 12-month development program. '
        'Voluntary turnover declined to 8.9%, well below the industry average of 13.2%.'
    )

    # 7
    doc.add_heading('Risk Management and Governance', level=2)
    doc.add_paragraph(
        'The Board of Directors strengthened our enterprise risk framework by establishing '
        'a dedicated Cybersecurity Committee. We completed SOC 2 Type II and ISO 27001 '
        'recertifications with zero major findings. Regulatory compliance costs decreased 12% '
        'through automation of monitoring and reporting processes. The company maintained '
        'an investment-grade credit rating of A- from Standard & Poor\'s.'
    )

    # 8
    doc.add_heading('Strategic Outlook for 2026', level=2)
    doc.add_paragraph(
        'Looking ahead, we expect revenue growth of 15-18% in fiscal year 2026, supported '
        'by continued momentum in cloud services and the ramp-up of our edge computing platform. '
        'Planned capital expenditures of $380 million will fund two new data centers in Frankfurt '
        'and Singapore. We remain focused on disciplined execution, customer-centric innovation, '
        'and responsible stewardship of stakeholder value.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
