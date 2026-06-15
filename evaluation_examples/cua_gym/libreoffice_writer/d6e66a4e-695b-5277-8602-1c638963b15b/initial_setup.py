"""
Initial Setup: Create a legal agreement document with Representations and Warranties section
where five items are listed as plain paragraphs (no bullet formatting).
Task ID: writer_legal_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_018'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('ASSET PURCHASE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_after = Pt(6)
    run = preamble.add_run(
        'This Asset Purchase Agreement (the "Agreement") is entered into as of '
        'March 15, 2025, by and between Meridian Technologies, Inc., a Delaware '
        'corporation ("Buyer"), and Cascade Digital Solutions, LLC, a California '
        'limited liability company ("Seller").'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Recitals ---
    doc.add_heading('RECITALS', level=1)

    recital_a = doc.add_paragraph()
    run = recital_a.add_run(
        'WHEREAS, Seller is engaged in the business of developing and licensing '
        'enterprise resource planning software solutions for mid-market manufacturing '
        'companies (the "Business");'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    recital_b = doc.add_paragraph()
    run = recital_b.add_run(
        'WHEREAS, Buyer desires to purchase from Seller, and Seller desires to sell '
        'to Buyer, substantially all of the assets used in connection with the Business, '
        'subject to the terms and conditions set forth herein;'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    recital_c = doc.add_paragraph()
    run = recital_c.add_run(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements '
        'hereinafter set forth and for other good and valuable consideration, the receipt '
        'and sufficiency of which are hereby acknowledged, the parties agree as follows:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Article I: Definitions ---
    doc.add_heading('ARTICLE I: DEFINITIONS', level=1)

    def_para = doc.add_paragraph()
    run = def_para.add_run(
        '"Acquired Assets" means all of the assets, properties, and rights of every '
        'kind and nature, whether real, personal, or mixed, tangible or intangible, '
        'owned or held by Seller and used in or related to the Business, including '
        'without limitation all intellectual property, customer contracts, equipment, '
        'inventory, and goodwill.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    def_para2 = doc.add_paragraph()
    run = def_para2.add_run(
        '"Purchase Price" means the aggregate consideration of Twelve Million Five '
        'Hundred Thousand Dollars ($12,500,000), payable in accordance with Section 3.1 '
        'of this Agreement.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Article II: Purchase and Sale ---
    doc.add_heading('ARTICLE II: PURCHASE AND SALE', level=1)

    sale_para = doc.add_paragraph()
    run = sale_para.add_run(
        'Subject to the terms and conditions of this Agreement, at the Closing, Seller '
        'shall sell, assign, transfer, convey, and deliver to Buyer, and Buyer shall '
        'purchase, acquire, and accept from Seller, the Acquired Assets, free and clear '
        'of all liens, encumbrances, and claims of any kind.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Article III: Representations and Warranties ---
    doc.add_heading('ARTICLE III: REPRESENTATIONS AND WARRANTIES', level=1)

    intro = doc.add_paragraph()
    run = intro.add_run(
        'Seller hereby represents and warrants to Buyer as follows:'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Five items as plain paragraphs (NO bullet formatting)
    items = [
        'Organization and Good Standing',
        'Authority and Capacity',
        'No Conflicts',
        'Compliance with Laws',
        'No Litigation',
    ]
    for item_text in items:
        p = doc.add_paragraph()
        run = p.add_run(item_text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # --- Article IV: Closing Conditions ---
    doc.add_heading('ARTICLE IV: CLOSING CONDITIONS', level=1)

    closing_para = doc.add_paragraph()
    run = closing_para.add_run(
        'The obligations of each party to consummate the transactions contemplated by '
        'this Agreement shall be subject to the fulfillment, at or prior to the Closing, '
        'of the conditions set forth in this Article IV, unless waived in writing by the '
        'party for whose benefit such condition exists.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # --- Article V: Indemnification ---
    doc.add_heading('ARTICLE V: INDEMNIFICATION', level=1)

    indem_para = doc.add_paragraph()
    run = indem_para.add_run(
        'Seller shall indemnify, defend, and hold harmless Buyer and its affiliates, '
        'officers, directors, employees, agents, and representatives from and against '
        'any and all losses, damages, liabilities, costs, and expenses arising out of '
        'or relating to any breach of any representation, warranty, covenant, or '
        'agreement made by Seller in this Agreement.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
