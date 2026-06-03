"""
Reward Script: Create bulleted list for Representations and Warranties section
Task ID: writer_legal_018
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.6): Each of the 5 items styled as 'List Bullet' (0.12 pts each)
  - Component 2 (0.2): All 5 items bulleted AND text content preserved (compound check)
  - Component 3 (0.2): Bullet items correctly placed under the right heading (requires >=3 bulleted)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_018'

EXPECTED_ITEMS = [
    'Organization and Good Standing',
    'Authority and Capacity',
    'No Conflicts',
    'Compliance with Laws',
    'No Litigation',
]

HEADING_TEXT = 'ARTICLE III: REPRESENTATIONS AND WARRANTIES'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the 'Representations and Warranties' section
    # Find paragraphs between the heading and the next heading
    section_paras = []
    in_section = False
    for para in doc.paragraphs:
        if HEADING_TEXT in para.text and para.style and 'Heading' in para.style.name:
            in_section = True
            continue
        if in_section:
            if para.style and 'Heading' in para.style.name:
                break  # reached next section
            section_paras.append(para)

    print(f"INFO: Found {len(section_paras)} paragraphs in Representations and Warranties section")

    # Component 1: Each of the 5 items styled as 'List Bullet' (0.12 pts each, total 0.6)
    # This is the core task-introduced change: style must change from Normal to List Bullet
    bullet_count = 0
    for item_text in EXPECTED_ITEMS:
        try:
            found = False
            for para in section_paras:
                if para.text.strip() == item_text:
                    found = True
                    style_name = para.style.name if para.style else 'None'
                    # Check for List Bullet style (could be 'List Bullet' or variants)
                    if 'List Bullet' in style_name:
                        print(f"PASS: '{item_text}' has style '{style_name}' (0.12 pts)")
                        total_score += 0.12
                        bullet_count += 1
                    else:
                        print(f"FAIL: '{item_text}' has style '{style_name}', expected 'List Bullet'")
                    break
            if not found:
                print(f"FAIL: '{item_text}' not found in section")
        except Exception as e:
            print(f"ERROR: Checking '{item_text}': {e}")

    # Component 2: All 5 items have List Bullet style AND text is preserved (0.2 pts)
    # This component awards points only when ALL items are correctly bulleted
    # Anchored to the task change: requires bullet_count == 5
    try:
        if bullet_count == 5:
            # Verify text content is also preserved (compound check)
            section_texts = [p.text.strip() for p in section_paras]
            items_found = sum(1 for item in EXPECTED_ITEMS if item in section_texts)
            if items_found == 5:
                print(f"PASS: All 5 items bulleted with correct text preserved (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: All bulleted but only {items_found}/5 item texts match")
        else:
            print(f"FAIL: Only {bullet_count}/5 items have List Bullet style")
    except Exception as e:
        print(f"ERROR: Checking complete bullet formatting: {e}")

    # Component 3: Bullet items correctly placed under the right heading (0.2 pts)
    # Anchored to task change: requires bullet_count >= 3
    try:
        heading_found = False
        for para in doc.paragraphs:
            if HEADING_TEXT in para.text and para.style and 'Heading' in para.style.name:
                heading_found = True
                break

        if heading_found and bullet_count >= 3:
            print(f"PASS: {bullet_count} bulleted items under '{HEADING_TEXT}' heading (0.2 pts)")
            total_score += 0.2
        elif heading_found and bullet_count > 0:
            print(f"PARTIAL: Only {bullet_count} bulleted items under heading")
        elif heading_found:
            print(f"FAIL: Heading found but no items have bullet style")
        else:
            print(f"FAIL: Heading '{HEADING_TEXT}' not found")
    except Exception as e:
        print(f"ERROR: Checking section structure: {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
