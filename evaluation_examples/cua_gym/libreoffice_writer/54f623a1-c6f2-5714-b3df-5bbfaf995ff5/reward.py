"""
Reward Script: Set French (France) language on second paragraph of page 2
Task ID: writer_fp_014
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): French paragraph has at least one run with lang fr-FR
  Component 2 (0.3): ALL runs in the French paragraph have lang fr-FR
  Component 3 (0.2): English paragraphs remain en-GB (no over-application)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_014'
FRENCH_TEXT_PREFIX = 'Cette etude examine'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_run_lang(run):
    """Extract the w:lang val attribute from a run's rPr, or None."""
    rPr = run.element.find(qn('w:rPr'))
    if rPr is not None:
        lang_elem = rPr.find(qn('w:lang'))
        if lang_elem is not None:
            return lang_elem.get(f'{{{W_NS}}}val')
    return None


def find_french_paragraph(doc):
    """Find the paragraph that starts with the known French text prefix."""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(FRENCH_TEXT_PREFIX):
            return i, para
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the French paragraph by its content
    french_idx, french_para = find_french_paragraph(doc)
    if french_para is None:
        print("CRITICAL: Could not find the French paragraph starting with "
              f"'{FRENCH_TEXT_PREFIX}'")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: French paragraph found at index {french_idx}")
    print(f"INFO: Text preview: {french_para.text[:80]}")

    runs = french_para.runs
    if not runs:
        print("FAIL: French paragraph has no runs")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: At least one run in the French paragraph has lang fr-FR (0.5 points)
    try:
        run_langs = [get_run_lang(r) for r in runs]
        any_french = any(lang and lang.lower().startswith('fr') for lang in run_langs)
        if any_french:
            print(f"PASS: Component 1 — At least one run has French language "
                  f"(langs: {run_langs}) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — No run has French language "
                  f"(langs: {run_langs})")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: ALL runs in the French paragraph have lang fr-FR (0.3 points)
    # This checks complete application, not just partial
    try:
        all_french = all(
            lang and lang.lower().startswith('fr')
            for lang in run_langs
        )
        if all_french:
            print(f"PASS: Component 2 — All {len(runs)} runs have French language "
                  f"(0.3 pts)")
            total_score += 0.3
        else:
            non_french = [(i, run_langs[i]) for i in range(len(run_langs))
                          if not (run_langs[i] and run_langs[i].lower().startswith('fr'))]
            print(f"FAIL: Component 2 — Not all runs are French. "
                  f"Non-French runs: {non_french}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: French paragraph is fr-FR AND English paragraphs remain en-GB (0.2 points)
    # This is a compound check: it only passes when the French paragraph is correctly
    # set AND English paragraphs are not over-applied. This ensures it FAILS on initial_env
    # (where the French paragraph is still en-GB).
    try:
        # Gate: French paragraph must already be detected as French (from Component 1)
        if not any_french:
            print("FAIL: Component 3 — French paragraph not set to French, "
                  "so compound check fails")
        else:
            english_ok = True
            checked = 0
            for i, para in enumerate(doc.paragraphs):
                if i == french_idx:
                    continue
                if not para.text.strip():
                    continue
                for run in para.runs:
                    lang = get_run_lang(run)
                    if lang and lang.lower().startswith('fr'):
                        print(f"FAIL: Component 3 — Paragraph {i} has French "
                              f"language but should be English: "
                              f"'{para.text[:40]}...' lang={lang}")
                        english_ok = False
                        break
                if not english_ok:
                    break
                checked += 1

            if english_ok and checked > 0:
                print(f"PASS: Component 3 — French para is fr-FR AND "
                      f"{checked} English paragraphs remain en-GB (0.2 pts)")
                total_score += 0.2
            elif checked == 0:
                print("FAIL: Component 3 — No English paragraphs found to verify")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 1)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
