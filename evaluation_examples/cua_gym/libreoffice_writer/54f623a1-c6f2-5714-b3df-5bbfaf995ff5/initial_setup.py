"""
Initial Setup: Bilingual report with French paragraph not yet marked as French
Task ID: writer_fp_014
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_014'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_lang(run, lang_code):
    """Set the w:lang val and bidi attributes on a run."""
    rPr = run._element.get_or_add_rPr()
    lang_elem = rPr.find(qn('w:lang'))
    if lang_elem is None:
        lang_elem = rPr.makeelement(qn('w:lang'), {})
        rPr.append(lang_elem)
    lang_elem.set(qn('w:val'), lang_code)


def add_paragraph_with_lang(doc, text, lang_code, style=None, bold=False, font_size=None, alignment=None):
    """Add a paragraph with all runs set to a specific language."""
    para = doc.add_paragraph(style=style)
    if alignment:
        para.paragraph_format.alignment = alignment
    run = para.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    run.font.name = 'Liberation Serif'
    set_run_lang(run, lang_code)
    return para


def create_initial():
    doc = Document()

    # Set default language to en-GB
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
        if rPrDefault is not None:
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is not None:
                lang_elem = rPr.find(qn('w:lang'))
                if lang_elem is None:
                    lang_elem = rPr.makeelement(qn('w:lang'), {})
                    rPr.append(lang_elem)
                lang_elem.set(qn('w:val'), 'en-GB')

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    EN = 'en-GB'

    # === PAGE 1 ===

    # Title
    add_paragraph_with_lang(doc, 'Impact of Climate Change on Marine Biodiversity',
                            EN, bold=True, font_size=18,
                            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    # Subtitle
    add_paragraph_with_lang(doc, 'A Comprehensive Review of Recent Research Findings',
                            EN, font_size=13,
                            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    # Blank line
    doc.add_paragraph()

    # Authors
    add_paragraph_with_lang(doc, 'Dr. Eleanor Whitfield, Dr. Jean-Pierre Moreau, Dr. Akiko Tanaka',
                            EN, font_size=11,
                            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    add_paragraph_with_lang(doc, 'Department of Marine Sciences, University of Bristol',
                            EN, font_size=11,
                            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    add_paragraph_with_lang(doc, 'Published: March 2025',
                            EN, font_size=11,
                            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    doc.add_paragraph()

    # Section 1
    add_paragraph_with_lang(doc, '1. Introduction', EN, bold=True, font_size=14)

    add_paragraph_with_lang(doc, (
        'Climate change represents one of the most significant threats to marine ecosystems '
        'worldwide. Rising ocean temperatures, acidification, and shifting current patterns '
        'are fundamentally altering the habitats upon which countless species depend. This '
        'report synthesises findings from over 150 peer-reviewed studies published between '
        '2020 and 2025, providing a comprehensive overview of the current state of knowledge '
        'regarding climate impacts on marine biodiversity.'
    ), EN, font_size=11)

    add_paragraph_with_lang(doc, (
        'The urgency of this review cannot be overstated. Recent data from the Intergovernmental '
        'Panel on Climate Change (IPCC) suggests that ocean warming is accelerating at a rate '
        'unprecedented in the geological record. The consequences for marine life are already '
        'visible: coral bleaching events have increased in frequency by 300% since the 1980s, '
        'and species migration patterns are shifting poleward at an average rate of 72 kilometres '
        'per decade.'
    ), EN, font_size=11)

    # Section 2
    add_paragraph_with_lang(doc, '2. Methodology', EN, bold=True, font_size=14)

    add_paragraph_with_lang(doc, (
        'Our methodological approach combined systematic literature review with meta-analysis '
        'techniques. We searched databases including Web of Science, Scopus, and PubMed using '
        'key terms related to marine biodiversity and climate change. Studies were evaluated '
        'against strict inclusion criteria: peer-reviewed publication, quantitative data on '
        'species abundance or distribution, and measurement of at least one climate variable.'
    ), EN, font_size=11)

    add_paragraph_with_lang(doc, (
        'Statistical analyses were performed using R version 4.3.2. Effect sizes were calculated '
        'using Hedges\u2019 g for continuous outcomes and odds ratios for categorical data. '
        'Heterogeneity was assessed using the I-squared statistic, with values above 75% '
        'indicating substantial heterogeneity. Publication bias was evaluated through funnel '
        'plots and Egger\u2019s regression test.'
    ), EN, font_size=11)

    add_paragraph_with_lang(doc, (
        'A total of 152 studies met our inclusion criteria, spanning 43 countries and covering '
        'marine environments from tropical coral reefs to Arctic ice shelves. The studies '
        'collectively examined over 1,200 marine species across 14 major taxonomic groups.'
    ), EN, font_size=11)

    # === PAGE 2 === (force page break)
    # Add a page break
    para_break = doc.add_paragraph()
    run_break = para_break.add_run()
    br = run_break._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run_break._element.append(br)

    # Section 3 - first paragraph on page 2
    add_paragraph_with_lang(doc, '3. Key Findings', EN, bold=True, font_size=14)

    # Second paragraph on page 2 - the French abstract
    # IMPORTANT: This is set to en-GB (NOT French) - this is the initial state
    # The task is to change this to French (France)
    add_paragraph_with_lang(doc, (
        'Cette etude examine les effets du changement climatique sur la biodiversite marine. '
        'Les resultats montrent une diminution significative de la richesse specifique dans '
        'les ecosystemes tropicaux, avec une perte moyenne de 14,2% des especes endemiques '
        'au cours de la derniere decennie. Les recifs coralliens sont particulierement '
        'vulnerables, avec 67% des sites etudies presentant des signes de blanchissement severe.'
    ), EN, font_size=11)

    # More English content
    add_paragraph_with_lang(doc, (
        'Our analysis reveals three primary mechanisms through which climate change affects '
        'marine biodiversity. First, thermal stress directly impacts species physiology, '
        'particularly in ectothermic organisms that cannot regulate their body temperature '
        'independently. Second, ocean acidification reduces the availability of carbonate '
        'ions essential for shell-forming organisms. Third, changes in ocean circulation '
        'patterns disrupt nutrient distribution and larval dispersal pathways.'
    ), EN, font_size=11)

    add_paragraph_with_lang(doc, (
        'The data demonstrate a clear latitudinal gradient in climate vulnerability. Tropical '
        'marine ecosystems, whilst harbouring the greatest species diversity, show the highest '
        'rates of biodiversity loss. Temperate regions exhibit more complex patterns, with some '
        'areas experiencing net increases in species richness as warm-water species expand their '
        'ranges northward.'
    ), EN, font_size=11)

    # Section 4
    add_paragraph_with_lang(doc, '4. Discussion', EN, bold=True, font_size=14)

    add_paragraph_with_lang(doc, (
        'These findings have profound implications for marine conservation policy. The rate '
        'of biodiversity loss observed in our meta-analysis exceeds the predictions of earlier '
        'models by approximately 40%. This discrepancy may be attributed to synergistic effects '
        'between climate stressors and other anthropogenic pressures, including overfishing, '
        'habitat destruction, and pollution from agricultural runoff.'
    ), EN, font_size=11)

    add_paragraph_with_lang(doc, (
        'Effective conservation strategies must therefore adopt an integrated approach that '
        'addresses multiple stressors simultaneously. Marine protected areas (MPAs) alone are '
        'insufficient; they must be complemented by broader measures to reduce greenhouse gas '
        'emissions and limit local anthropogenic impacts on marine habitats.'
    ), EN, font_size=11)

    # Section 5
    add_paragraph_with_lang(doc, '5. Conclusion', EN, bold=True, font_size=14)

    add_paragraph_with_lang(doc, (
        'This comprehensive review underscores the critical threat that climate change poses '
        'to marine biodiversity. Immediate and coordinated international action is required '
        'to mitigate the worst effects. Future research should prioritise long-term monitoring '
        'programmes and the development of predictive models that incorporate multiple climate '
        'variables and their interactions with anthropogenic stressors.'
    ), EN, font_size=11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
