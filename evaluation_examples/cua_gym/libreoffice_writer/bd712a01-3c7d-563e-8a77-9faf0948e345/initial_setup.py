"""
Initial Setup: Insert a Table of Contents at the beginning of this document based on the document headings.
Task ID: osworld_writer_toc_generation_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_toc_generation_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Document Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("Meridian Technologies Inc.")
    title_run.bold = True
    title_run.font.size = Pt(20)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle_para.add_run("Annual Business Report 2024")
    sub_run.italic = True
    sub_run.font.size = Pt(14)

    doc.add_paragraph()  # spacing

    # Section 1: Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Meridian Technologies Inc. was founded in 2010 with the mission to deliver "
        "innovative enterprise software solutions to mid-market and large enterprises "
        "across North America and Europe. Over the past fourteen years, we have grown "
        "from a team of twelve employees to a workforce of over 1,400 professionals "
        "operating across eight countries."
    )

    doc.add_heading("Company Overview", level=2)
    doc.add_paragraph(
        "Headquartered in Austin, Texas, Meridian Technologies specializes in cloud-based "
        "enterprise resource planning (ERP), customer relationship management (CRM), and "
        "data analytics platforms. Our flagship product, MeridianERP Pro, currently serves "
        "over 3,200 active enterprise clients worldwide, with a customer retention rate of 94%."
    )

    doc.add_heading("Report Scope and Objectives", level=2)
    doc.add_paragraph(
        "This annual report covers fiscal year 2024, encompassing the period from January 1, 2024 "
        "to December 31, 2024. The report presents a comprehensive review of business performance, "
        "financial results, operational developments, and strategic initiatives undertaken during "
        "the reporting period. It also outlines our vision and planned investments for 2025."
    )

    # Section 2: Market Analysis
    doc.add_heading("Market Analysis", level=1)
    doc.add_paragraph(
        "The enterprise software market continued its robust expansion in 2024, driven by accelerating "
        "digital transformation initiatives across all major industry verticals. Total addressable market "
        "for cloud ERP solutions reached approximately $48.2 billion globally, representing a year-over-year "
        "growth rate of 13.7% compared to 2023."
    )

    doc.add_heading("Competitive Landscape", level=2)
    doc.add_paragraph(
        "The competitive environment remains dynamic, with established players such as SAP, Oracle, and "
        "Microsoft Dynamics maintaining significant market share. However, Meridian Technologies has "
        "successfully differentiated through its AI-powered analytics layer, flexible deployment model, "
        "and superior customer success program. During 2024, we achieved a net promoter score of 72, "
        "placing us among the top three vendors in our segment."
    )

    doc.add_heading("Industry Trends and Opportunities", level=2)
    doc.add_paragraph(
        "Key market trends observed during 2024 include the rapid adoption of generative AI in business "
        "workflows, the continued migration from on-premise to hybrid cloud architectures, and increased "
        "regulatory requirements driving demand for compliance management features. Meridian's product "
        "roadmap is well-positioned to capitalize on each of these trends, as detailed in the Future "
        "Outlook section of this report."
    )

    # Section 3: Financial Performance
    doc.add_heading("Financial Performance", level=1)
    doc.add_paragraph(
        "Fiscal year 2024 marked another year of strong financial performance for Meridian Technologies. "
        "Total revenue reached $312.4 million, representing a 19.6% increase over the prior year's "
        "revenue of $261.2 million. This growth was primarily driven by expansion of our subscription "
        "revenue base and successful cross-selling of premium analytics modules to existing customers."
    )

    doc.add_heading("Revenue Breakdown", level=2)
    doc.add_paragraph(
        "Subscription revenue accounted for $248.7 million (79.6% of total revenue), up from $198.3 million "
        "in 2023. Professional services revenue contributed $41.8 million (13.4%), while hardware and "
        "perpetual license revenue represented the remaining $21.9 million (7.0%). The shift toward "
        "subscription-based revenue continues to improve our recurring revenue predictability and "
        "customer lifetime value metrics."
    )

    doc.add_heading("Profitability and Cash Flow", level=2)
    doc.add_paragraph(
        "Gross profit margin improved to 68.3% in 2024, compared to 65.7% in 2023, reflecting operating "
        "leverage in our cloud infrastructure and improved service delivery efficiency. Adjusted EBITDA "
        "reached $74.9 million (24.0% margin), and free cash flow generation was $58.2 million, up 34% "
        "year-over-year. We ended the fiscal year with a cash and equivalents balance of $143.6 million "
        "and no long-term debt."
    )

    # Section 4: Operations Review
    doc.add_heading("Operations Review", level=1)
    doc.add_paragraph(
        "The Operations division made substantial progress in 2024 across product development, customer "
        "support, and infrastructure reliability. We successfully launched MeridianERP Pro version 7.0 in "
        "March 2024, which introduced the AI Assistant feature, redesigned dashboard system, and enhanced "
        "integration marketplace with over 280 pre-built connectors."
    )

    doc.add_heading("Product Development", level=2)
    doc.add_paragraph(
        "The engineering team, now comprising 420 engineers across three development hubs in Austin, "
        "Krakow, and Singapore, delivered 6 major product releases and 42 minor updates during 2024. "
        "R&D investment totaled $52.1 million (16.7% of revenue), with a particular focus on machine "
        "learning capabilities, mobile-first user experience improvements, and enhanced security features "
        "including SOC 2 Type II and ISO 27001 certification renewals."
    )

    doc.add_heading("Customer Success and Support", level=2)
    doc.add_paragraph(
        "Our Customer Success organization expanded to 180 professionals dedicated to onboarding, "
        "training, and ongoing adoption support. Average time-to-value for new enterprise clients was "
        "reduced from 94 days to 67 days through the implementation of our accelerated deployment "
        "methodology. Customer support ticket resolution time improved by 22%, and our self-service "
        "knowledge base now handles 61% of tier-1 support inquiries without human intervention."
    )

    # Section 5: Future Outlook
    doc.add_heading("Future Outlook", level=1)
    doc.add_paragraph(
        "Looking ahead to fiscal year 2025, Meridian Technologies is well-positioned for continued "
        "growth and innovation. We have established four strategic priorities that will guide our "
        "investments and operational focus throughout the coming year."
    )

    doc.add_heading("Strategic Priorities for 2025", level=2)
    doc.add_paragraph(
        "Our primary strategic initiatives for 2025 include: (1) launching MeridianAI Suite, our "
        "next-generation platform with embedded generative AI workflows across all modules; (2) "
        "expanding into the Asia-Pacific market with dedicated operations in Japan, Australia, and "
        "South Korea; (3) acquiring complementary technology capabilities in the supply chain "
        "optimization space; and (4) achieving carbon-neutral operations across all facilities "
        "by Q3 2025."
    )

    doc.add_heading("Financial Guidance", level=2)
    doc.add_paragraph(
        "For fiscal year 2025, management is providing the following guidance: total revenue in the "
        "range of $368 million to $380 million (18-22% growth), adjusted EBITDA margin of 25-27%, "
        "and free cash flow generation of $70-80 million. Capital expenditures are expected to be "
        "approximately $18 million, primarily related to the Asia-Pacific expansion and data center "
        "infrastructure upgrades. These projections are based on current market conditions and "
        "may be subject to revision as the year progresses."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
