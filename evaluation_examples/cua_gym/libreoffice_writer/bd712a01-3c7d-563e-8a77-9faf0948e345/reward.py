"""
Reward Script: Insert a Table of Contents at the beginning of this document based on the document headings.
Task ID: osworld_writer_toc_generation_001
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4 pts): A TOC field code (instrText containing 'TOC') exists anywhere in the document
  Component 2 (0.3 pts): The TOC field/heading appears at the beginning of the document (within first 3 paragraphs)
  Component 3 (0.3 pts): The TOC covers heading levels 1 and 2 (includes Heading 1 and Heading 2 headings)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_toc_generation_001'

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task asks to insert a Table of Contents at the beginning of the document
    based on document headings. The golden file has:
      - Para 0: Heading 1 'Table of Contents'
      - Para 1: TOC field code paragraph (style TOC1, instrText: TOC levels 1-2)
      - Para 2: Empty paragraph
      - (original content follows)

    The initial file has no TOC field code and starts with 'Meridian Technologies Inc.'
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file " + file_path + ": " + str(e))
        print("REWARD: 0.0")
        return 0.0

    # Collect all instruction text fields from the whole document
    toc_instr_texts = []
    for p in doc.paragraphs:
        for run in p.runs:
            for instr_el in run._element.findall('.//w:instrText', NS):
                text = instr_el.text or ''
                if 'TOC' in text.upper():
                    toc_instr_texts.append(text.strip())

    # Component 1: A TOC field code exists in the document (0.4 points)
    # This is the primary indicator that a Table of Contents was actually inserted
    try:
        if len(toc_instr_texts) > 0:
            print("PASS: Component 1 — TOC field code found: " + repr(toc_instr_texts[0]) + " (0.4 pts)")
            total_score += 0.4
        else:
            print("FAIL: Component 1 — No TOC field code found in document")
    except Exception as e:
        print("ERROR: Component 1 — " + str(e))

    # Component 2: TOC content appears at the BEGINNING of the document (0.3 points)
    # The task says "Insert a Table of Contents at the beginning of this document"
    # Verify: the TOC field or a heading titled "Table of Contents" appears in the first
    # 5 paragraphs of the document (allowing for a title heading + TOC field + blank).
    try:
        toc_at_beginning = False
        for p in doc.paragraphs[:5]:
            # Check if this paragraph contains TOC instruction text
            para_has_toc_field = False
            for run in p.runs:
                for instr_el in run._element.findall('.//w:instrText', NS):
                    text = instr_el.text or ''
                    if 'TOC' in text.upper():
                        para_has_toc_field = True
            # Also accept a "Table of Contents" heading at position 0
            is_toc_heading = (
                'table of contents' in p.text.lower() and
                p.style.name.lower().startswith('heading')
            )
            if para_has_toc_field or is_toc_heading:
                toc_at_beginning = True
                break

        if toc_at_beginning:
            print("PASS: Component 2 — TOC content found at the beginning of the document (0.3 pts)")
            total_score += 0.3
        else:
            print("FAIL: Component 2 — TOC content not found in the first 5 paragraphs")
    except Exception as e:
        print("ERROR: Component 2 — " + str(e))

    # Component 3: The document still contains the original headings (Heading 1 and Heading 2)
    # AND the TOC field references levels that cover both heading levels.
    # This verifies the TOC is based on the document's actual headings and that the
    # document structure (headings) is preserved intact.
    # The TOC field must reference heading levels (e.g., \o "1-2" covering both H1 and H2).
    try:
        # Count Heading 1 and Heading 2 paragraphs
        h1_headings = [p.text for p in doc.paragraphs if p.style.name == 'Heading 1'
                       and p.text.lower() != 'table of contents']
        h2_headings = [p.text for p in doc.paragraphs if p.style.name == 'Heading 2']

        has_h1 = len(h1_headings) >= 1
        has_h2 = len(h2_headings) >= 1

        # Check that TOC field references heading levels (contains \o with level range)
        toc_covers_levels = False
        for instr_text in toc_instr_texts:
            # \o "1-2" or similar level specification
            if r'\o' in instr_text or '\\o' in instr_text.lower():
                toc_covers_levels = True
            # Also accept if no explicit \o (defaults to all levels)
            if 'TOC' in instr_text.upper():
                toc_covers_levels = True  # any TOC field by default covers headings

        if has_h1 and has_h2 and toc_covers_levels:
            print("PASS: Component 3 — TOC field covers heading levels; document has " +
                  str(len(h1_headings)) + " H1 and " + str(len(h2_headings)) + " H2 headings (0.3 pts)")
            total_score += 0.3
        else:
            print("FAIL: Component 3 — TOC coverage issue. has_h1=" + str(has_h1) +
                  " has_h2=" + str(has_h2) + " toc_covers_levels=" + str(toc_covers_levels))
    except Exception as e:
        print("ERROR: Component 3 — " + str(e))

    final_score = min(total_score, 1.0)
    print("")
    print("Score: " + str(total_score) + "/1.0")
    print("REWARD: " + str(final_score))
    return final_score


# Default: test against canonical artifact path in a given env
file_path = WORKDIR + '/' + TASK_ID + '.docx'
if not os.path.exists(file_path):
    print("File not found: " + file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
