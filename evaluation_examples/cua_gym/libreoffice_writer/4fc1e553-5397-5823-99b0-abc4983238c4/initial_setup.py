"""
Initial Setup: Single-column Writer document with title and body text
Task ID: writer_fs_050
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_050'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard letter size with 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title_para = doc.add_heading('Annual Report 2024', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 1
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(10)
    run1 = p1.add_run(
        'The fiscal year 2024 marked a significant turning point for Meridian Technologies. '
        'Revenue grew by 18.3% year-over-year, reaching $247.5 million, driven primarily by '
        'strong demand in the cloud infrastructure division and a 32% increase in enterprise '
        'contract renewals. Our strategic investments in artificial intelligence and machine '
        'learning capabilities have begun to yield measurable returns, with the AI Solutions '
        'segment contributing $41.2 million in its first full year of operation.'
    )
    run1.font.size = Pt(11)
    run1.font.name = 'Calibri'

    # Paragraph 2
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(
        'Our workforce expanded to 3,850 employees across 14 global offices, reflecting a '
        'net addition of 620 positions. The majority of new hires joined the engineering and '
        'product development teams, underscoring our commitment to innovation. Employee '
        'satisfaction scores reached an all-time high of 4.3 out of 5.0, attributed to the '
        'introduction of flexible work arrangements and enhanced professional development '
        'programs launched in the second quarter.'
    )
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'

    # Paragraph 3
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(10)
    run3 = p3.add_run(
        'The Research and Development division allocated $38.7 million toward next-generation '
        'products, representing 15.6% of total revenue. Key breakthroughs included the launch '
        'of the Quantum Edge platform, which processes real-time analytics 40 times faster than '
        'our previous generation systems. Patent filings increased by 27%, with 43 new patents '
        'granted across cloud computing, cybersecurity, and data orchestration technologies.'
    )
    run3.font.size = Pt(11)
    run3.font.name = 'Calibri'

    # Paragraph 4
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(10)
    run4 = p4.add_run(
        'Sustainability remained a core pillar of our corporate strategy. Carbon emissions were '
        'reduced by 22% compared to the 2023 baseline, exceeding our initial target of 15%. '
        'All data centers now operate on 100% renewable energy sources, and the company achieved '
        'ISO 14001 certification for environmental management. Community engagement initiatives, '
        'including the TechBridge mentorship program, benefited over 12,000 students in '
        'underserved communities.'
    )
    run4.font.size = Pt(11)
    run4.font.name = 'Calibri'

    # Paragraph 5
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(10)
    run5 = p5.add_run(
        'Looking ahead to 2025, Meridian Technologies is positioned to capitalize on expanding '
        'market opportunities in edge computing and autonomous systems. The board has approved '
        'a $52 million capital expenditure plan focused on infrastructure modernization and '
        'geographic expansion into Southeast Asia and Latin America. We remain confident that '
        'our disciplined approach to growth, combined with a culture of relentless innovation, '
        'will deliver sustained value for our shareholders, partners, and employees.'
    )
    run5.font.size = Pt(11)
    run5.font.name = 'Calibri'

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
