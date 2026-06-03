"""
Reward Script: Convert five agenda items into a default bulleted list
Task ID: writer_lec_001
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): Proportional credit for agenda items with 'List Bullet' style
  Component 2 (0.4): All 5 items are 'List Bullet' AND original text is preserved
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_001'

# Expected agenda item texts (from task context)
EXPECTED_ITEMS = [
    'Review Q3 results',
    'Discuss hiring plan',
    'Budget allocation for Q4',
    'Product roadmap updates',
    'Team building event planning',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 6 paragraphs (1 heading + 5 items)
    if len(doc.paragraphs) < 6:
        print(f"FAIL: Expected at least 6 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Identify the 5 agenda item paragraphs (paragraphs 1-5, after the heading)
    agenda_paras = doc.paragraphs[1:6]

    # Component 1: Proportional credit for agenda items with 'List Bullet' style (0.6 points)
    # Each item with 'List Bullet' earns 0.12 points (0.6 / 5)
    try:
        bullet_count = 0
        for i, para in enumerate(agenda_paras):
            style_name = para.style.name if para.style else 'None'
            if style_name == 'List Bullet':
                bullet_count += 1
                print(f"PASS: Para {i+1} has 'List Bullet' style (text: '{para.text}')")
            else:
                print(f"FAIL: Para {i+1} has style '{style_name}', expected 'List Bullet' (text: '{para.text}')")

        component1_score = (bullet_count / 5) * 0.6
        if bullet_count > 0:
            print(f"PASS: Component 1 -- {bullet_count}/5 items have 'List Bullet' style ({component1_score:.2f} pts)")
            total_score += component1_score
        else:
            print(f"FAIL: Component 1 -- No items have 'List Bullet' style")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: All 5 items are 'List Bullet' AND original text is preserved (0.4 points)
    # This component requires BOTH conditions: all bullets applied AND text unchanged
    try:
        all_bullet = (bullet_count == 5)
        mismatched = [
            i for i, para in enumerate(agenda_paras)
            if para.text.strip() != EXPECTED_ITEMS[i]
        ]
        for i in mismatched:
            print(f"FAIL: Component 2 -- Para {i+1} text mismatch: expected '{EXPECTED_ITEMS[i]}', found '{agenda_paras[i].text.strip()}'")

        if all_bullet and len(mismatched) == 0:
            print(f"PASS: Component 2 -- All 5 items are bulleted with correct text (0.4 pts)")
            total_score += 0.4
        elif not all_bullet:
            print(f"FAIL: Component 2 -- Not all items are bulleted ({bullet_count}/5)")
        elif len(mismatched) > 0:
            print(f"FAIL: Component 2 -- Text content was modified")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
