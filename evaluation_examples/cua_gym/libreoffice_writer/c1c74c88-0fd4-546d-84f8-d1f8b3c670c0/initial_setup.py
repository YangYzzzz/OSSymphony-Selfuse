"""
Initial Setup: Create a Writer document with five plain-text meeting agenda items.
Task ID: writer_lec_001
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    doc.add_heading('Team Meeting Agenda', level=1)

    # Five plain-text agenda items -- each on its own paragraph, NO bullet style
    agenda_items = [
        'Review Q3 results',
        'Discuss hiring plan',
        'Budget allocation for Q4',
        'Product roadmap updates',
        'Team building event planning',
    ]

    for item in agenda_items:
        doc.add_paragraph(item)  # default "Normal" style -- no bullets

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Open in LibreOffice Writer for the GUI agent
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
