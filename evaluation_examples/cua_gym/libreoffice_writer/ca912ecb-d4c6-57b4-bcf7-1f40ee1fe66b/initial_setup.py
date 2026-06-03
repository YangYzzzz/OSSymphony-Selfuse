"""
Initial Setup: Internal memo document with 5 pages, no header configured.
Task ID: writer_fs_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # === PAGE 1 — Memo Title & Introduction ===
    title = doc.add_heading('Greenfield Technologies Inc.', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Internal Memorandum')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph('')  # spacer

    # Memo metadata
    meta_items = [
        ('TO:', 'All Department Heads and Senior Managers'),
        ('FROM:', 'Rebecca Torres, Chief Operations Officer'),
        ('DATE:', 'March 28, 2026'),
        ('RE:', 'Q2 2026 Strategic Initiatives and Operational Changes'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        r1 = p.add_run(label + '  ')
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    # Horizontal line (border bottom on a paragraph)
    hr = doc.add_paragraph()
    hr_fmt = hr.paragraph_format
    hr_fmt.space_after = Pt(6)
    pBdr = hr._element.get_or_add_pPr()
    from docx.oxml.ns import qn
    bdr = pBdr.makeelement(qn('w:pBdr'), {})
    bottom = bdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:space'): '1',
        qn('w:color'): '2E74B5',
    })
    bdr.append(bottom)
    pBdr.append(bdr)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        'This memo outlines the key strategic initiatives and operational changes '
        'planned for Q2 2026. All department heads are expected to review the details '
        'and prepare implementation timelines for their respective teams by April 15, 2026.'
    )
    run.font.size = Pt(11)

    # === PAGE 2 — Strategic Priorities ===
    doc.add_page_break()

    h2 = doc.add_heading('1. Strategic Priorities for Q2 2026', level=1)

    priorities = [
        ('1.1 Market Expansion into Southeast Asia',
         'Following the successful pilot program in Singapore during Q1, we will be '
         'expanding operations to Vietnam, Thailand, and the Philippines. Regional '
         'Director Anh Nguyen will coordinate the rollout with local partners. '
         'Initial capital allocation is $3.2M with projected ROI of 18% by end of fiscal year.'),
        ('1.2 Product Line Consolidation',
         'The product management team under James Whitfield has completed the portfolio '
         'analysis. We will sunset the Legacy Connect platform (effective June 30) and '
         'migrate all 2,847 active users to the CloudBridge Pro suite. Migration support '
         'staff will be temporarily expanded by 12 FTEs from April through July.'),
        ('1.3 AI Integration Roadmap',
         'Chief Technology Officer Priya Sharma will present the full AI integration '
         'roadmap at the April 8 leadership meeting. Key highlights include automated '
         'customer support triage (reducing response time by 40%), predictive inventory '
         'management for our logistics division, and intelligent document processing '
         'for the legal department.'),
    ]

    for sub_title, body_text in priorities:
        h3 = doc.add_heading(sub_title, level=2)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(body_text)
        run.font.size = Pt(11)

    # === PAGE 3 — Operational Changes ===
    doc.add_page_break()

    doc.add_heading('2. Operational Changes', level=1)

    doc.add_heading('2.1 Revised Hybrid Work Policy', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'Effective May 1, 2026, the hybrid work policy will be updated as follows. '
        'All employees in Bands 4-7 are required to be on-site a minimum of three days '
        'per week (Tuesday, Wednesday, and one flexible day). Remote-only exceptions '
        'require VP-level approval. The facilities team has redesigned Building C, Floor 3 '
        'as a collaborative workspace with hot-desking capability for 180 employees.'
    )
    run.font.size = Pt(11)

    doc.add_heading('2.2 Budget Reallocation', level=2)

    # Add a table for budget details
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Department', 'Q1 Budget ($K)', 'Q2 Budget ($K)', 'Change (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    budget_data = [
        ['Engineering', '1,250', '1,480', '+18.4%'],
        ['Marketing', '890', '720', '-19.1%'],
        ['Operations', '650', '685', '+5.4%'],
        ['Human Resources', '420', '395', '-6.0%'],
        ['Research & Dev', '1,100', '1,350', '+22.7%'],
    ]
    for r_idx, row_data in enumerate(budget_data, 1):
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph('')  # spacer

    doc.add_heading('2.3 Vendor Consolidation', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'The procurement team has completed the vendor audit. We will consolidate from '
        '47 active vendors to 28 preferred partners. CFO Daniel Park has negotiated '
        'volume discounts with key suppliers that are projected to save $890K annually. '
        'All departments must transition to approved vendors by May 31.'
    )
    run.font.size = Pt(11)

    # === PAGE 4 — HR & Compliance Updates ===
    doc.add_page_break()

    doc.add_heading('3. Human Resources & Compliance', level=1)

    doc.add_heading('3.1 New Hiring Targets', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'Total Q2 hiring target: 45 positions across the organization. Engineering '
        'will receive 22 headcount (including 8 senior ML engineers and 6 platform '
        'engineers). Marketing receives 8 positions focused on digital campaigns '
        'and market research. Operations adds 10 logistics coordinators for the '
        'Southeast Asia expansion. HR adds 5 recruiters on 6-month contracts to '
        'support the hiring surge.'
    )
    run.font.size = Pt(11)

    doc.add_heading('3.2 Mandatory Training Programs', level=2)
    bullets = [
        'Cybersecurity Awareness Refresher (all employees, due April 30)',
        'Anti-Harassment and DEI Workshop (managers, due May 15)',
        'Data Privacy Compliance (GDPR/CCPA) Update (customer-facing roles, due May 31)',
        'Emergency Response Protocol Training (facilities and security, due April 22)',
        'New Employee Onboarding Redesign (HR facilitators, pilot April 14)',
    ]
    for bullet in bullets:
        bp = doc.add_paragraph(bullet, style='List Bullet')
        for run in bp.runs:
            run.font.size = Pt(11)

    doc.add_heading('3.3 Benefits Enrollment Changes', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'Open enrollment for the updated benefits package runs from April 1-15. '
        'Notable changes include the addition of a mental health stipend ($1,200/year), '
        'expanded parental leave (16 weeks for all parents), and the introduction of a '
        'student loan repayment assistance program (up to $5,000/year for eligible employees). '
        'HR Director Monica Reeves will host information sessions on April 3 and April 7.'
    )
    run.font.size = Pt(11)

    # === PAGE 5 — Action Items & Timeline ===
    doc.add_page_break()

    doc.add_heading('4. Action Items and Timeline', level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        'Each department head is responsible for the following deliverables. '
        'Progress reports are due bi-weekly starting April 14.'
    )
    run.font.size = Pt(11)

    action_items = [
        ('April 4', 'Department heads acknowledge receipt of this memo and confirm '
         'review with direct reports.'),
        ('April 15', 'Submit Q2 implementation timelines to COO office.'),
        ('April 22', 'Complete vendor transition plans (procurement liaison assigned).'),
        ('April 30', 'Cybersecurity training completion verification.'),
        ('May 1', 'New hybrid work policy takes effect.'),
        ('May 15', 'Mid-quarter progress review meeting (all department heads).'),
        ('May 31', 'Vendor consolidation complete. GDPR/CCPA training deadline.'),
        ('June 15', 'Southeast Asia expansion Phase 1 launch.'),
        ('June 30', 'Legacy Connect platform sunset. Full CloudBridge Pro migration complete.'),
    ]

    # Action items as numbered list
    for date, desc in action_items:
        p = doc.add_paragraph(style='List Number')
        r1 = p.add_run(date + ' — ')
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)

    doc.add_paragraph('')  # spacer

    closing = doc.add_paragraph()
    run = closing.add_run(
        'Please do not hesitate to reach out to my office if you have questions or '
        'require clarification on any of the above items. I look forward to a productive quarter.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph('')

    sign = doc.add_paragraph()
    run = sign.add_run('Rebecca Torres')
    run.bold = True
    run.font.size = Pt(11)

    sign2 = doc.add_paragraph()
    run = sign2.add_run('Chief Operations Officer')
    run.font.size = Pt(11)

    sign3 = doc.add_paragraph()
    run = sign3.add_run('Greenfield Technologies Inc.')
    run.font.size = Pt(11)
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
