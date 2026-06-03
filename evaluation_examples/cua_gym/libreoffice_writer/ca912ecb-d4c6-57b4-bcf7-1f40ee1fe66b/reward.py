"""
Reward Script: Set up header to not appear on first page but appear on subsequent pages
Task ID: writer_fs_076
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): titlePg enabled (different first page header)
  Component 2 (0.25): First page header is empty (no header on page 1)
  Component 3 (0.30): Default header contains 'Internal Memo - Confidential'
  Component 4 (0.20): Default header is right-aligned
"""

import os

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_076'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section
    if len(doc.sections) == 0:
        print("CRITICAL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: titlePg element is present (different first page header enabled) (0.25 points)
    try:
        sect_pr = section._sectPr
        title_pg = sect_pr.find(qn('w:titlePg'))
        if title_pg is not None:
            print(f"PASS: Component 1 — titlePg element present (different first page enabled) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — titlePg element not found; different first page header not enabled")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: First page header is empty AND titlePg is enabled (0.25 points)
    # This compound check ensures we only score when "different first page" is active
    # AND the first page header has no text (i.e., header suppressed on page 1).
    try:
        sect_pr2 = section._sectPr
        title_pg2 = sect_pr2.find(qn('w:titlePg'))
        first_hdr = section.first_page_header
        first_hdr_text = ""
        if first_hdr and first_hdr.paragraphs:
            first_hdr_text = "".join(p.text for p in first_hdr.paragraphs).strip()
        if title_pg2 is not None and first_hdr_text == "":
            print(f"PASS: Component 2 — titlePg enabled AND first page header is empty (0.25 pts)")
            total_score += 0.25
        else:
            if title_pg2 is None:
                print(f"FAIL: Component 2 — titlePg not enabled, cannot verify first page header suppression")
            else:
                print(f"FAIL: Component 2 — First page header has text: {repr(first_hdr_text)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Default header contains 'Internal Memo - Confidential' (0.30 points)
    try:
        header = section.header
        header_text = ""
        if header and header.paragraphs:
            header_text = "".join(p.text for p in header.paragraphs).strip()
        if 'Internal Memo - Confidential' in header_text:
            print(f"PASS: Component 3 — Default header contains 'Internal Memo - Confidential' (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 3 — Default header text: {repr(header_text)}, expected 'Internal Memo - Confidential'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Default header text is right-aligned (0.20 points)
    try:
        header = section.header
        if header and header.paragraphs:
            # Find the paragraph with the header text and check its alignment
            right_aligned_found = any(
                p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT
                for p in header.paragraphs
                if p.text.strip()
            )
            if right_aligned_found:
                print(f"PASS: Component 4 — Header text is right-aligned (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — Header text is not right-aligned")
        else:
            print(f"FAIL: Component 4 — No header paragraphs to check alignment")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved GUI state
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
