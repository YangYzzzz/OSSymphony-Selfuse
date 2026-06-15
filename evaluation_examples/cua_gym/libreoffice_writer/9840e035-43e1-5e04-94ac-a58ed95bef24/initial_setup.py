"""
Initial Setup: Writer document with images but no auto-captioning configured
Task ID: writer_tech_077
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_077'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_sample_image(path, width, height, color, label):
    """Create a simple colored image with a label for realistic content."""
    img = Image.new('RGB', (width, height), color)
    img.save(path)


def create_initial():
    # Create sample images for the document
    img1_path = f'{WORKDIR}/project_architecture.png'
    img2_path = f'{WORKDIR}/performance_metrics.png'
    img3_path = f'{WORKDIR}/deployment_pipeline.png'

    create_sample_image(img1_path, 640, 400, (45, 85, 130), 'Architecture')
    create_sample_image(img2_path, 640, 350, (70, 130, 80), 'Metrics')
    create_sample_image(img3_path, 640, 380, (140, 70, 50), 'Pipeline')

    doc = Document()

    # --- Title ---
    title = doc.add_heading('Cloud Migration Technical Report', level=0)

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=1)
    intro = doc.add_paragraph(
        'This report outlines the technical architecture and migration strategy for '
        'transitioning the Meridian Analytics Platform from on-premises infrastructure '
        'to Amazon Web Services (AWS). The migration encompasses 47 microservices, '
        '12 databases, and 3 data processing pipelines currently serving approximately '
        '2.3 million monthly active users across North America and Europe.'
    )

    doc.add_paragraph(
        'The project timeline spans Q2 2025 through Q4 2025, with a phased approach '
        'that minimizes service disruption while maximizing cost optimization opportunities. '
        'Key stakeholders include the Platform Engineering team, led by Director Sarah Nakamura, '
        'and the DevOps division under VP Marcus Chen.'
    )

    # --- Architecture section with image ---
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        'The target architecture follows a containerized microservices pattern deployed '
        'on Amazon EKS (Elastic Kubernetes Service). Each service communicates through '
        'an internal service mesh powered by AWS App Mesh, with external traffic routed '
        'through Application Load Balancers across three availability zones.'
    )

    # Insert first image (no caption - that's the task)
    doc.add_picture(img1_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        'The architecture separates concerns into four primary layers: ingestion, '
        'processing, storage, and presentation. The ingestion layer handles approximately '
        '850,000 events per second during peak hours, utilizing Amazon Kinesis Data Streams '
        'with a 7-day retention period for replay capability.'
    )

    # --- Performance section with image ---
    doc.add_heading('3. Performance Benchmarks', level=1)
    doc.add_paragraph(
        'Load testing conducted in March 2025 demonstrated that the containerized '
        'architecture handles 3.2x the throughput of the legacy monolithic deployment. '
        'P99 latency decreased from 340ms to 89ms for the primary API endpoints, '
        'while memory utilization dropped by 41% due to more efficient resource allocation.'
    )

    # Insert second image
    doc.add_picture(img2_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        'Database performance showed significant improvement after migrating from '
        'self-managed PostgreSQL 14 to Amazon Aurora PostgreSQL. Read replica lag '
        'decreased from an average of 120ms to under 10ms, and automated failover '
        'reduced recovery time from 45 seconds to approximately 3 seconds.'
    )

    # --- Deployment section with image ---
    doc.add_heading('4. CI/CD Pipeline and Deployment', level=1)
    doc.add_paragraph(
        'The deployment pipeline leverages AWS CodePipeline integrated with GitHub Actions '
        'for continuous integration. Each merge to the main branch triggers automated builds, '
        'unit tests (target: 92% coverage), integration tests against staging environments, '
        'and canary deployments to production with automated rollback capabilities.'
    )

    # Insert third image
    doc.add_picture(img3_path, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        'Blue-green deployments are used for database schema changes, ensuring zero-downtime '
        'migrations. The team has established a deployment cadence of twice weekly for feature '
        'releases, with hotfix capability within a 30-minute window from commit to production.'
    )

    # --- Summary ---
    doc.add_heading('5. Cost Analysis', level=1)
    doc.add_paragraph(
        'Projected annual infrastructure costs after migration total $1.28 million, '
        'representing a 34% reduction from the current $1.94 million on-premises expenditure. '
        'The largest savings come from right-sizing compute resources (estimated $380K annually) '
        'and eliminating datacenter lease obligations ($290K annually).'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up temp images
    for p in [img1_path, img2_path, img3_path]:
        if os.path.exists(p):
            os.remove(p)

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
