"""
Reward Script: Certificate of Completion in LibreOffice Writer
Task ID: writer_rd_091
Domain: libreoffice_writer
Scoring:
  - Component 1: Double-line page border with gold color (0.15)
  - Component 2: 3.0 cm margins (0.10)
  - Component 3: "Certificate of Completion" title - centered, bold, ~28pt, small caps (0.20)
  - Component 4: "This certifies that" centered ~14pt (0.10)
  - Component 5: Recipient name placeholder - underscores, script font, ~22pt (0.10)
  - Component 6: Course name "Advanced Data Analytics Course" bold ~16pt (0.15)
  - Component 7: Signature lines with Instructor/Director labels (0.10)
  - Component 8: Date fields at bottom (0.10)
"""

import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_091'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify certificate of completion task with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have content (initial doc is blank)
    if len(doc.paragraphs) == 0:
        print("PRECONDITION FAIL: Document has no paragraphs - appears to be blank/initial state")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: Double-line page border with gold color CC9900 (0.15 points)
    try:
        pgBorders = section._sectPr.findall('.//w:pgBorders', ns)
        if pgBorders:
            pb = pgBorders[0]
            sides = ['top', 'left', 'bottom', 'right']
            border_ok = True
            for side in sides:
                elem = pb.find(f'w:{side}', ns)
                if elem is None:
                    border_ok = False
                    print(f"FAIL: Component 1 — missing {side} border element")
                    break
                val = elem.get(f'{{{ns["w"]}}}val', '')
                color = elem.get(f'{{{ns["w"]}}}color', '').upper()
                if 'double' not in val.lower():
                    border_ok = False
                    print(f"FAIL: Component 1 — {side} border style is '{val}', expected 'double'")
                    break
                if color != 'CC9900':
                    border_ok = False
                    print(f"FAIL: Component 1 — {side} border color is '{color}', expected 'CC9900'")
                    break
            if border_ok:
                print("PASS: Component 1 — Double-line page border with gold #CC9900 (0.15 pts)")
                total_score += 0.15
        else:
            print("FAIL: Component 1 — No page borders found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 3.0 cm margins on all sides (0.10 points)
    try:
        # 3.0 cm = 1080000 EMU (360000 EMU per cm). Allow ±0.2 cm tolerance
        target_margin = 1080000
        tolerance = 72000  # 0.2 cm
        margins = {
            'left': section.left_margin,
            'right': section.right_margin,
            'top': section.top_margin,
            'bottom': section.bottom_margin,
        }
        margins_ok = True
        for name, val in margins.items():
            if val is None or abs(val - target_margin) > tolerance:
                margins_ok = False
                actual_cm = val / 360000 if val else 0
                print(f"FAIL: Component 2 — {name} margin is {actual_cm:.2f} cm, expected ~3.0 cm")
                break
        if margins_ok:
            print("PASS: Component 2 — All margins ~3.0 cm (0.10 pts)")
            total_score += 0.10
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Helper: collect all paragraph texts for searching
    all_paras = doc.paragraphs
    para_texts = [p.text.strip() for p in all_paras]

    # Component 3: "Certificate of Completion" title - centered, bold, ~28pt, small caps (0.20 points)
    try:
        found_title = False
        for para in all_paras:
            if 'certificate' in para.text.lower() and 'completion' in para.text.lower():
                # Check centered
                align = para.paragraph_format.alignment
                is_centered = align == WD_PARAGRAPH_ALIGNMENT.CENTER
                # Check bold and size on runs
                has_bold = False
                has_approx_size = False
                has_small_caps = False
                for run in para.runs:
                    if run.font.bold:
                        has_bold = True
                    if run.font.size and 24 <= run.font.size.pt <= 36:
                        has_approx_size = True
                    # Check small caps via XML
                    rPr = run._element.find('.//w:rPr', ns)
                    if rPr is not None:
                        sc = rPr.find('w:smallCaps', ns)
                        if sc is not None:
                            sc_val = sc.get(f'{{{ns["w"]}}}val', 'true')
                            if sc_val.lower() in ('true', '1', 'on', ''):
                                has_small_caps = True

                sub_score = 0.0
                if is_centered and has_bold:
                    sub_score += 0.10
                if has_approx_size:
                    sub_score += 0.05
                if has_small_caps:
                    sub_score += 0.05

                if sub_score > 0:
                    found_title = True
                    total_score += sub_score
                    print(f"PASS: Component 3 — 'Certificate of Completion' found (centered={is_centered}, bold={has_bold}, size_ok={has_approx_size}, smallCaps={has_small_caps}) ({sub_score} pts)")
                else:
                    print(f"FAIL: Component 3 — Title found but formatting wrong (centered={is_centered}, bold={has_bold}, size_ok={has_approx_size}, smallCaps={has_small_caps})")
                break
        if not found_title:
            print("FAIL: Component 3 — 'Certificate of Completion' text not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: "This certifies that" centered ~14pt (0.10 points)
    try:
        found = False
        for para in all_paras:
            if 'this certifies that' in para.text.lower():
                is_centered = para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                size_ok = False
                for run in para.runs:
                    if run.font.size and 12 <= run.font.size.pt <= 18:
                        size_ok = True
                if is_centered:
                    found = True
                    total_score += 0.10
                    print(f"PASS: Component 4 — 'This certifies that' centered, size_ok={size_ok} (0.10 pts)")
                else:
                    print(f"FAIL: Component 4 — 'This certifies that' found but not centered")
                break
        if not found:
            print("FAIL: Component 4 — 'This certifies that' text not found or not centered")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Recipient name placeholder - underscores, script-like font, ~22pt (0.10 points)
    try:
        found = False
        for para in all_paras:
            text = para.text.strip()
            # Look for a paragraph that is primarily underscores (name placeholder)
            if text.count('_') >= 10 and 'date' not in text.lower() and 'instructor' not in text.lower() and 'director' not in text.lower():
                # Check if this is the name placeholder (not signature lines)
                # Name placeholder should be around 22pt with script font
                has_script_font = False
                has_large_size = False
                for run in para.runs:
                    if run.font.name and 'script' in run.font.name.lower():
                        has_script_font = True
                    if run.font.size and 18 <= run.font.size.pt <= 28:
                        has_large_size = True

                # Only count if it looks like the name placeholder (large size differentiates from sig lines)
                if has_large_size or has_script_font:
                    found = True
                    sub_score = 0.05 if has_large_size else 0.0
                    sub_score += 0.05 if has_script_font else 0.0
                    total_score += sub_score
                    print(f"PASS: Component 5 — Name placeholder found (script_font={has_script_font}, size_ok={has_large_size}) ({sub_score} pts)")
                    break
        if not found:
            print("FAIL: Component 5 — Recipient name placeholder not found with script font / ~22pt")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: "has successfully completed" + "Advanced Data Analytics Course" bold ~16pt (0.15 points)
    try:
        found_completed = False
        found_course = False
        for para in all_paras:
            lower_text = para.text.lower().strip()
            if 'has successfully completed' in lower_text:
                found_completed = True
            if 'advanced data analytics course' in lower_text:
                # Check bold
                is_bold = any(run.font.bold for run in para.runs)
                size_ok = any(run.font.size and 14 <= run.font.size.pt <= 20 for run in para.runs)
                if is_bold:
                    found_course = True
                    print(f"  Course title: bold={is_bold}, size_ok={size_ok}")

        sub_score = 0.0
        if found_completed:
            sub_score += 0.05
        if found_course:
            sub_score += 0.10
        if sub_score > 0:
            total_score += sub_score
            print(f"PASS: Component 6 — 'has successfully completed'={found_completed}, course_bold={found_course} ({sub_score} pts)")
        else:
            print(f"FAIL: Component 6 — Missing 'has successfully completed' or 'Advanced Data Analytics Course'")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Signature lines with Instructor/Director labels (0.10 points)
    try:
        found_instructor = False
        found_director = False
        found_sig_lines = False
        for para in all_paras:
            lower_text = para.text.lower()
            if 'instructor' in lower_text:
                found_instructor = True
            if 'director' in lower_text:
                found_director = True
            # Signature lines: underscores not in date context, in the bottom portion
            if para.text.count('_') >= 8 and 'date' not in lower_text:
                # Could be name placeholder or signature - check if instructor/director nearby
                pass

        # Check for signature underscores in paragraphs near instructor/director labels
        for para in all_paras:
            if para.text.count('_') >= 8 and ('instructor' in para.text.lower() or 'director' in para.text.lower()):
                found_sig_lines = True
            # Also check if a separate paragraph has sig-line underscores
            # (they might be on the same line if using tabs/spaces)
        # Also look for a line with multiple underscore groups (two signature areas)
        for para in all_paras:
            text = para.text.strip()
            underscore_groups = [g for g in text.split() if g.count('_') >= 5]
            if len(underscore_groups) >= 2 and 'date' not in text.lower():
                found_sig_lines = True

        sub_score = 0.0
        if found_instructor:
            sub_score += 0.03
        if found_director:
            sub_score += 0.03
        if found_sig_lines:
            sub_score += 0.04
        if sub_score > 0:
            total_score += sub_score
            print(f"PASS: Component 7 — Instructor={found_instructor}, Director={found_director}, sig_lines={found_sig_lines} ({sub_score} pts)")
        else:
            print("FAIL: Component 7 — No signature lines with Instructor/Director labels found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: Date fields at bottom (0.10 points)
    try:
        found_date = False
        date_count = 0
        for para in all_paras:
            lower_text = para.text.lower()
            if 'date' in lower_text and '_' in para.text:
                found_date = True
                # Count how many "date" references
                date_count += lower_text.count('date')

        if found_date and date_count >= 2:
            total_score += 0.10
            print(f"PASS: Component 8 — Date fields found ({date_count} date references) (0.10 pts)")
        elif found_date:
            total_score += 0.05
            print(f"PARTIAL: Component 8 — Only {date_count} date field(s) found (0.05 pts)")
        else:
            print("FAIL: Component 8 — No date fields found")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
