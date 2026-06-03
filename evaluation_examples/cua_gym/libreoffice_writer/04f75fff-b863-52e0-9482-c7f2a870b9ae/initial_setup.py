"""
Initial Setup: Product catalog document with single-column layout
Task ID: writer_fs_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: A4, 2.5 cm left/right margins ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # Single column (default) -- no column XML needed

    # --- Title ---
    title = doc.add_heading('Meridian Home & Living', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Product Catalog  |  Spring/Summer 2025')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    # --- Introduction ---
    doc.add_heading('Welcome to Our Collection', level=1)

    intro = doc.add_paragraph()
    intro.add_run(
        'At Meridian Home & Living, we believe that every space tells a story. '
        'Our curated selection of home furnishings and decor pieces is designed '
        'to help you create environments that are both beautiful and functional. '
        'From handcrafted artisan furniture to modern minimalist accessories, '
        'each item in our catalog has been carefully chosen for quality, style, '
        'and sustainability.'
    )

    intro2 = doc.add_paragraph()
    intro2.add_run(
        'This season, we are thrilled to introduce over 40 new products across '
        'our core categories. Whether you are refreshing a single room or '
        'redesigning your entire home, our Spring/Summer 2025 collection offers '
        'something for every taste and budget.'
    )

    # --- Living Room ---
    doc.add_heading('Living Room Essentials', level=1)

    doc.add_heading('Sofas & Seating', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Aurora Modular Sofa')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(
        '\nA versatile, L-shaped sectional available in charcoal velvet or '
        'sand linen. Features removable cushion covers and hidden storage '
        'compartments. Dimensions: 280 cm x 180 cm x 85 cm.'
    )
    p.add_run('\nPrice: $2,450.00  |  SKU: MLH-SOF-1024')

    p2 = doc.add_paragraph()
    run2 = p2.add_run('Bentley Wingback Armchair')
    run2.bold = True
    run2.font.size = Pt(11)
    p2.add_run(
        '\nClassic wingback silhouette reimagined with contemporary fabric '
        'choices including emerald green, burnt sienna, and midnight blue. '
        'Solid oak frame with brass-tipped legs.'
    )
    p2.add_run('\nPrice: $899.00  |  SKU: MLH-CHR-0587')

    doc.add_heading('Coffee Tables & Side Tables', level=2)

    p3 = doc.add_paragraph()
    run3 = p3.add_run('Cascade Marble Coffee Table')
    run3.bold = True
    run3.font.size = Pt(11)
    p3.add_run(
        '\nItalian Carrara marble top with a brushed gold geometric base. '
        'A statement piece that anchors any living space. '
        'Dimensions: 120 cm x 60 cm x 42 cm.'
    )
    p3.add_run('\nPrice: $1,680.00  |  SKU: MLH-TBL-0293')

    p4 = doc.add_paragraph()
    run4 = p4.add_run('Elm Nesting Side Tables (Set of 3)')
    run4.bold = True
    run4.font.size = Pt(11)
    p4.add_run(
        '\nSolid elm wood with natural grain finish. Stackable design saves '
        'space while offering flexible surface options for drinks, books, '
        'and decorative objects.'
    )
    p4.add_run('\nPrice: $420.00  |  SKU: MLH-TBL-0314')

    # --- Bedroom ---
    doc.add_heading('Bedroom Collection', level=1)

    doc.add_heading('Beds & Frames', level=2)

    p5 = doc.add_paragraph()
    run5 = p5.add_run('Serene Platform Bed Frame')
    run5.bold = True
    run5.font.size = Pt(11)
    p5.add_run(
        '\nMinimalist walnut platform bed with integrated headboard and '
        'floating nightstands. Available in queen and king sizes. '
        'No box spring required.'
    )
    p5.add_run('\nPrice: $1,890.00 (Queen) / $2,190.00 (King)  |  SKU: MLH-BED-0745')

    doc.add_heading('Bedding & Textiles', level=2)

    p6 = doc.add_paragraph()
    run6 = p6.add_run('Luxe Egyptian Cotton Sheet Set')
    run6.bold = True
    run6.font.size = Pt(11)
    p6.add_run(
        '\n800 thread count, 100% long-staple Egyptian cotton. Includes '
        'fitted sheet, flat sheet, and two pillowcases. Colors: ivory, '
        'dove grey, blush pink, ocean blue.'
    )
    p6.add_run('\nPrice: $289.00  |  SKU: MLH-BDG-0412')

    # --- Kitchen & Dining ---
    doc.add_heading('Kitchen & Dining', level=1)

    p7 = doc.add_paragraph()
    run7 = p7.add_run('Farmhouse Extending Dining Table')
    run7.bold = True
    run7.font.size = Pt(11)
    p7.add_run(
        '\nReclaimed pine dining table with butterfly leaf extension. '
        'Seats 6 to 8 comfortably. Hand-distressed finish gives each '
        'table a unique character. Dimensions: 180-240 cm x 95 cm x 76 cm.'
    )
    p7.add_run('\nPrice: $1,350.00  |  SKU: MLH-DIN-0158')

    p8 = doc.add_paragraph()
    run8 = p8.add_run('Copenhagen Dining Chair (Pair)')
    run8.bold = True
    run8.font.size = Pt(11)
    p8.add_run(
        '\nScandinavian-inspired design with curved beech wood back and '
        'woven paper cord seat. Sold in pairs. Stackable for easy storage.'
    )
    p8.add_run('\nPrice: $540.00 per pair  |  SKU: MLH-DIN-0173')

    # --- Ordering Info ---
    doc.add_heading('Ordering Information', level=1)

    order_info = doc.add_paragraph()
    order_info.add_run(
        'All items are available for order through our website at '
        'www.meridianhome.com or by calling our customer service line '
        'at +1 (800) 555-0192. Standard delivery is complimentary on orders '
        'over $500. White-glove delivery and assembly services are available '
        'for furniture items at an additional fee of $149.'
    )

    returns = doc.add_paragraph()
    returns.add_run(
        'We offer a 30-day satisfaction guarantee on all products. Items '
        'must be returned in original packaging and in unused condition. '
        'Custom-order and personalized items are non-refundable. Please '
        'refer to our full return policy at www.meridianhome.com/returns.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
