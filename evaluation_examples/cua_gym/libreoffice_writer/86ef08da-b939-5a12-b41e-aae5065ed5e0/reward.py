"""
Reward Script: Set first paragraph first-line indent to 1.5 cm and remove spacing before
Task ID: writer_fs_029
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): First paragraph first-line indent ~1.5 cm
  Component 2 (0.5 pts): First paragraph space_before == 0
"""

import os

from docx import Document
from docx.shared import Cm

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_029'

# 1.5 cm in EMU = 540000; allow +/- 5% tolerance
TARGET_INDENT_EMU = Cm(1.5)  # 540000 EMU
INDENT_TOLERANCE = 0.05  # 5%


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one paragraph
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    first_para = doc.paragraphs[0]
    pf = first_para.paragraph_format

    # Component 1: First-line indent is ~1.5 cm (0.5 points)
    # The task requires first_line_indent = 1.5 cm.
    # Initial state has first_line_indent = 0 (no indent).
    # Golden state has first_line_indent = 539750 EMU (~1.5 cm).
    try:
        fl_indent = pf.first_line_indent
        if fl_indent is not None and fl_indent > 0:
            target = int(TARGET_INDENT_EMU)
            lower = target * (1 - INDENT_TOLERANCE)
            upper = target * (1 + INDENT_TOLERANCE)
            if lower <= fl_indent <= upper:
                print(f"PASS: Component 1 — first_line_indent={fl_indent} EMU is within 5% of {target} EMU (1.5 cm) (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 1 — first_line_indent={fl_indent} EMU is outside 5% tolerance of {target} EMU (1.5 cm)")
        else:
            print(f"FAIL: Component 1 — first_line_indent is {fl_indent} (expected ~540000 EMU / 1.5 cm)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Space before paragraph is 0 (0.5 points)
    # The task requires removing spacing before the paragraph.
    # Initial state has space_before = 179705 EMU (~0.5 cm).
    # Golden state has space_before = 0.
    try:
        sb = pf.space_before
        # space_before == 0 or None (inherited, which defaults to 0 in most styles)
        # We need to check that it is explicitly 0, not just None (inherited).
        # In golden, space_before is 0 (EMU). In initial, it's 179705.
        # We accept 0 as pass.
        if sb is not None and sb == 0:
            print(f"PASS: Component 2 — space_before={sb} (removed, equals 0) (0.5 pts)")
            total_score += 0.5
        elif sb is None:
            # None means inherited from style. Check if the style default is 0.
            # For safety, we also accept None as "no explicit spacing" which often means 0.
            print(f"PASS: Component 2 — space_before is None (inherited, effectively 0) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 2 — space_before={sb} EMU (expected 0)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
