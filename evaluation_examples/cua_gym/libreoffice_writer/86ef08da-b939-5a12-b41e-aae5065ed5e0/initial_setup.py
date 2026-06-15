"""
Initial Setup: Set first paragraph indent and remove spacing before
Task ID: writer_fs_029
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins for a professional letter
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- First paragraph: body opening with NO indentation and 0.5 cm spacing before ---
    para1 = doc.add_paragraph(
        "We are pleased to inform you that your application for the Senior Project Manager "
        "position at Meridian Consulting Group has been reviewed by our hiring committee. "
        "After careful consideration of your qualifications and experience, we would like "
        "to invite you to the next stage of our selection process."
    )
    para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para1.paragraph_format.first_line_indent = Cm(0)  # No indentation
    para1.paragraph_format.space_before = Cm(0.5)     # 0.5 cm spacing above
    para1.paragraph_format.space_after = Pt(6)
    for run in para1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Second paragraph ---
    para2 = doc.add_paragraph(
        "Your extensive background in managing cross-functional teams and delivering "
        "complex infrastructure projects aligns well with our current organizational needs. "
        "The committee was particularly impressed by your leadership during the Eastbridge "
        "Development Initiative, which resulted in a 23% improvement in delivery timelines."
    )
    para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para2.paragraph_format.space_before = Pt(0)
    para2.paragraph_format.space_after = Pt(6)
    for run in para2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Third paragraph ---
    para3 = doc.add_paragraph(
        "We would like to schedule a panel interview at our downtown office located at "
        "450 Commerce Boulevard, Suite 1200, on Thursday, April 17, 2026, at 10:00 AM. "
        "The interview will last approximately 90 minutes and will include a brief "
        "presentation component where you will have the opportunity to share your vision "
        "for optimizing project workflows."
    )
    para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para3.paragraph_format.space_before = Pt(0)
    para3.paragraph_format.space_after = Pt(6)
    for run in para3.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Fourth paragraph ---
    para4 = doc.add_paragraph(
        "Please confirm your availability by responding to this letter no later than "
        "Friday, April 11, 2026. If you have any questions regarding the interview format "
        "or require any accommodations, do not hesitate to contact our Human Resources "
        "department at hr@meridianconsulting.com or by phone at (415) 782-3900."
    )
    para4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para4.paragraph_format.space_before = Pt(0)
    para4.paragraph_format.space_after = Pt(6)
    for run in para4.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Closing paragraph ---
    para5 = doc.add_paragraph("We look forward to meeting you and discussing how your expertise "
                              "can contribute to the continued success of Meridian Consulting Group.")
    para5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para5.paragraph_format.space_before = Pt(0)
    para5.paragraph_format.space_after = Pt(12)
    for run in para5.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # --- Signature block ---
    para6 = doc.add_paragraph("Sincerely,")
    para6.paragraph_format.space_before = Pt(24)
    para6.paragraph_format.space_after = Pt(0)
    for run in para6.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    para7 = doc.add_paragraph("Elena Vasquez")
    para7.paragraph_format.space_before = Pt(24)
    para7.paragraph_format.space_after = Pt(0)
    run7 = para7.runs[0]
    run7.bold = True
    run7.font.name = "Times New Roman"
    run7.font.size = Pt(12)

    para8 = doc.add_paragraph("Director of Talent Acquisition")
    para8.paragraph_format.space_before = Pt(0)
    para8.paragraph_format.space_after = Pt(0)
    for run in para8.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    para9 = doc.add_paragraph("Meridian Consulting Group")
    para9.paragraph_format.space_before = Pt(0)
    para9.paragraph_format.space_after = Pt(0)
    for run in para9.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
