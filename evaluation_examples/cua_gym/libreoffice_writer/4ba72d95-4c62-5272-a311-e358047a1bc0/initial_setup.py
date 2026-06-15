"""
Initial Setup: Insert CONFIDENTIAL at beginning of contract document
Task ID: writer_edit_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'contract_draft'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Heading: Service Agreement ---
    heading = doc.add_heading('Service Agreement', level=1)

    # --- Introduction ---
    intro = doc.add_paragraph(
        'This Service Agreement ("Agreement") is entered into as of March 1, 2025, '
        'by and between Meridian Consulting Group LLC ("Service Provider"), a Delaware '
        'limited liability company, and Brightfield Industries Inc. ("Client"), a California '
        'corporation.'
    )

    # --- Section 1 ---
    doc.add_heading('1. Services', level=2)
    doc.add_paragraph(
        'The Service Provider agrees to provide the following professional services to the Client: '
        'strategic management consulting, operational process analysis, and organizational restructuring '
        'advisory services. The scope of work shall be defined in detail in Exhibit A attached hereto.'
    )

    # --- Section 2 ---
    doc.add_heading('2. Term', level=2)
    doc.add_paragraph(
        'This Agreement shall commence on March 1, 2025, and shall continue in full force and effect '
        'until February 28, 2026, unless earlier terminated in accordance with the provisions herein. '
        'Either party may renew this Agreement for successive one-year terms upon written notice.'
    )

    # --- Section 3 ---
    doc.add_heading('3. Compensation', level=2)
    doc.add_paragraph(
        'In consideration of the services rendered, the Client shall pay the Service Provider a monthly '
        'retainer fee of $12,500 USD, due on the first business day of each calendar month. '
        'Additional project-based fees shall be agreed upon in writing prior to commencement of any '
        'out-of-scope work.'
    )

    # --- Section 4 ---
    doc.add_heading('4. Confidentiality', level=2)
    doc.add_paragraph(
        'Both parties agree to maintain the strict confidentiality of all proprietary information, '
        'trade secrets, and business data exchanged during the course of this Agreement. Neither party '
        'shall disclose such information to any third party without the prior written consent of the '
        'disclosing party. This obligation shall survive the termination of this Agreement for a period '
        'of five (5) years.'
    )

    # --- Section 5 ---
    doc.add_heading('5. Intellectual Property', level=2)
    doc.add_paragraph(
        'All deliverables, reports, and work product created by the Service Provider specifically for '
        'the Client under this Agreement shall become the sole property of the Client upon receipt of '
        'full payment. The Service Provider retains the right to use general methodologies and frameworks '
        'developed independently of this engagement.'
    )

    # --- Section 6 ---
    doc.add_heading('6. Limitation of Liability', level=2)
    doc.add_paragraph(
        'In no event shall the Service Provider be liable for any indirect, incidental, consequential, '
        'or punitive damages arising out of or relating to this Agreement, even if the Service Provider '
        'has been advised of the possibility of such damages. The total cumulative liability of the '
        'Service Provider shall not exceed the total fees paid in the three months preceding the claim.'
    )

    # --- Section 7 ---
    doc.add_heading('7. Termination', level=2)
    doc.add_paragraph(
        'Either party may terminate this Agreement upon thirty (30) days written notice to the other '
        'party. In the event of material breach by either party, the non-breaching party may terminate '
        'immediately upon written notice specifying the nature of the breach.'
    )

    # --- Section 8 ---
    doc.add_heading('8. Governing Law', level=2)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws of the State '
        'of Delaware, without regard to its conflict of law provisions. Any disputes arising hereunder '
        'shall be resolved through binding arbitration in Wilmington, Delaware.'
    )

    # --- Signature Block ---
    doc.add_heading('Signatures', level=2)
    doc.add_paragraph(
        'IN WITNESS WHEREOF, the parties have executed this Service Agreement as of the date first '
        'written above.'
    )

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.cell(0, 0).text = 'SERVICE PROVIDER'
    sig_table.cell(0, 1).text = 'CLIENT'
    sig_table.cell(1, 0).text = 'Meridian Consulting Group LLC'
    sig_table.cell(1, 1).text = 'Brightfield Industries Inc.'
    sig_table.cell(2, 0).text = 'By: ________________________\nName: Jonathan R. Caldwell\nTitle: Managing Partner\nDate: ___________________'
    sig_table.cell(2, 1).text = 'By: ________________________\nName: Diana L. Westbrook\nTitle: Chief Executive Officer\nDate: ___________________'

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
