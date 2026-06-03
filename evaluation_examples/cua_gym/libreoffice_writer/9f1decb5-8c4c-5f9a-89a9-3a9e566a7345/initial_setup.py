"""
Initial Setup: Screenplay scene document - unformatted initial state
Task ID: writer_creative_012
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_creative_012'
OUTPUT = f'{WORKDIR}/screenplay_scene.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font for the document (Arial 12pt, unformatted)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Set normal style to Arial 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = None
    font.italic = None

    def add_para(text, font_name='Arial', font_size=12, bold=False, italic=False,
                 alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, left_indent=None, right_indent=None):
        """Add a paragraph with specified formatting."""
        para = doc.add_paragraph()
        para.paragraph_format.alignment = alignment
        if left_indent is not None:
            para.paragraph_format.left_indent = left_indent
        if right_indent is not None:
            para.paragraph_format.right_indent = right_indent
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        return para

    # --- Title ---
    add_para(
        'INT. COFFEE SHOP - MORNING',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- JACK (first appearance) ---
    add_para(
        'JACK',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_para(
        "Look, I've been thinking about this all week. Maybe it's time we just... let things be.",
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- Stage direction ---
    add_para(
        '[Jack enters the room slowly]',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- ELENA (first appearance) ---
    add_para(
        'ELENA',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_para(
        "Let things be? Jack, do you hear yourself? Everything we worked for is sitting in that folder.",
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_para(
        "And you want to just walk away.",
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- Stage direction ---
    add_para(
        '[She looks out the window]',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- NARRATOR ---
    add_para(
        'NARRATOR',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_para(
        "The coffee shop hummed with the quiet noise of morning routines — cups clinking, chairs scraping, and the soft murmur of conversations that would never be remembered. Outside, the city moved without them.",
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- JACK (second appearance) ---
    add_para(
        'JACK',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_para(
        "It's not walking away. It's knowing when something's done.",
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- Stage direction ---
    add_para(
        '[Long pause]',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    # --- ELENA (second appearance) ---
    add_para(
        'ELENA',
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_para(
        "Then I guess we see things differently.",
        font_name='Arial', font_size=12,
        bold=False, italic=False,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also place on Desktop for convenience
    import shutil
    desktop_path = f'{WORKDIR}/Desktop/screenplay_scene.docx'
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)
    shutil.copy(OUTPUT, desktop_path)
    print(f'Copied to Desktop: {desktop_path}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
