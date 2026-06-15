"""
Initial Setup: Create a project status report document with a table.
Task ID: writer_biz_051
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_051'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Q1 2025 Project Status Report', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    doc.add_paragraph(
        'The following table summarizes the current status of all active projects '
        'under the Digital Transformation initiative. Project managers are expected '
        'to update their respective entries by the 15th of each month.'
    )

    doc.add_paragraph('')  # spacer

    # Project status data
    headers = ['Task', 'Owner', 'Due Date', 'Status']
    data = [
        ['Website Redesign', 'Sarah Chen', '2025-02-28', 'Complete'],
        ['CRM Migration', 'Marcus Johnson', '2025-03-15', 'In Progress'],
        ['Data Warehouse Setup', 'Priya Patel', '2025-01-31', 'Complete'],
        ['Mobile App Launch', 'David Kim', '2025-04-30', 'In Progress'],
        ['Security Audit', 'Elena Rodriguez', '2025-03-01', 'Delayed'],
        ['API Gateway Deployment', 'James Okafor', '2025-02-15', 'Complete'],
        ['Cloud Cost Optimization', 'Anika Gupta', '2025-05-15', 'In Progress'],
        ['Legacy System Decommission', 'Robert Tanaka', '2025-03-31', 'Delayed'],
        ['Employee Onboarding Portal', 'Lisa Moreau', '2025-04-15', 'In Progress'],
        ['Vendor Payment Integration', 'Carlos Reyes', '2025-02-28', 'Complete'],
    ]

    # Create table
    table = doc.add_table(rows=1 + len(data), cols=4)
    table.style = 'Table Grid'

    # Header row - bold
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Data rows - no background colors (that's the agent's task)
    for i, row_data in enumerate(data, 1):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(11)

    # Set reasonable column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(1.8)
        row.cells[2].width = Inches(1.2)
        row.cells[3].width = Inches(1.2)

    doc.add_paragraph('')  # spacer

    doc.add_paragraph(
        'Note: Projects marked as Delayed require escalation to the steering committee '
        'for resource reallocation review.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
