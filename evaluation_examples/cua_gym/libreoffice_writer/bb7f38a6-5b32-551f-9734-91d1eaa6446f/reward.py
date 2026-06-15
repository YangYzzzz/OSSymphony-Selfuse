"""
Reward Script: Conditional-style background formatting in Status column
Task ID: writer_biz_051
Domain: libreoffice_writer
Scoring:
  Component 1 (0.40): 'Complete' cells have green background (#C6EFCE)
  Component 2 (0.35): 'In Progress' cells have yellow background (#FFEB9C)
  Component 3 (0.25): 'Delayed' cells have red background (#FFC7CE)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_051'

# Expected color mapping (status text -> fill color hex, uppercase no #)
STATUS_COLORS = {
    'Complete': 'C6EFCE',
    'In Progress': 'FFEB9C',
    'Delayed': 'FFC7CE',
}


def get_cell_fill(cell):
    """Extract the w:shd fill color from a table cell, or None."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    if fill and fill.lower() != 'auto':
        return fill.upper()
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table must have header row with 'Status' column
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    if 'Status' not in headers:
        print(f"FAIL: No 'Status' column in table headers: {headers}")
        print("REWARD: 0.0")
        return 0.0

    status_col_idx = headers.index('Status')
    data_rows = list(table.rows)[1:]  # skip header

    # Build mapping of status values to their rows
    complete_rows = []
    in_progress_rows = []
    delayed_rows = []

    for ri, row in enumerate(data_rows, start=1):
        status_text = row.cells[status_col_idx].text.strip()
        if status_text == 'Complete':
            complete_rows.append((ri, row))
        elif status_text == 'In Progress':
            in_progress_rows.append((ri, row))
        elif status_text == 'Delayed':
            delayed_rows.append((ri, row))

    # Component 1: 'Complete' cells have green background #C6EFCE (0.40 points)
    try:
        expected_color = STATUS_COLORS['Complete']
        if len(complete_rows) == 0:
            print("FAIL: Component 1 -- No 'Complete' rows found")
        else:
            matches = 0
            for ri, row in complete_rows:
                fill = get_cell_fill(row.cells[status_col_idx])
                if fill == expected_color:
                    matches += 1
                else:
                    print(f"FAIL: Component 1 -- Row {ri} 'Complete' expected fill {expected_color}, found {fill}")
            fraction = matches / len(complete_rows)
            points = round(0.40 * fraction, 4)
            if matches == len(complete_rows):
                print(f"PASS: Component 1 -- All {matches} 'Complete' cells have green background #{expected_color} (0.40 pts)")
            else:
                print(f"PARTIAL: Component 1 -- {matches}/{len(complete_rows)} 'Complete' cells correct ({points} pts)")
            if points > 0:
                total_score += points
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: 'In Progress' cells have yellow background #FFEB9C (0.35 points)
    try:
        expected_color = STATUS_COLORS['In Progress']
        if len(in_progress_rows) == 0:
            print("FAIL: Component 2 -- No 'In Progress' rows found")
        else:
            matches = 0
            for ri, row in in_progress_rows:
                fill = get_cell_fill(row.cells[status_col_idx])
                if fill == expected_color:
                    matches += 1
                else:
                    print(f"FAIL: Component 2 -- Row {ri} 'In Progress' expected fill {expected_color}, found {fill}")
            fraction = matches / len(in_progress_rows)
            points = round(0.35 * fraction, 4)
            if matches == len(in_progress_rows):
                print(f"PASS: Component 2 -- All {matches} 'In Progress' cells have yellow background #{expected_color} (0.35 pts)")
            else:
                print(f"PARTIAL: Component 2 -- {matches}/{len(in_progress_rows)} 'In Progress' cells correct ({points} pts)")
            if points > 0:
                total_score += points
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: 'Delayed' cells have red background #FFC7CE (0.25 points)
    try:
        expected_color = STATUS_COLORS['Delayed']
        if len(delayed_rows) == 0:
            print("FAIL: Component 3 -- No 'Delayed' rows found")
        else:
            matches = 0
            for ri, row in delayed_rows:
                fill = get_cell_fill(row.cells[status_col_idx])
                if fill == expected_color:
                    matches += 1
                else:
                    print(f"FAIL: Component 3 -- Row {ri} 'Delayed' expected fill {expected_color}, found {fill}")
            fraction = matches / len(delayed_rows)
            points = round(0.25 * fraction, 4)
            if matches == len(delayed_rows):
                print(f"PASS: Component 3 -- All {matches} 'Delayed' cells have red background #{expected_color} (0.25 pts)")
            else:
                print(f"PARTIAL: Component 3 -- {matches}/{len(delayed_rows)} 'Delayed' cells correct ({points} pts)")
            if points > 0:
                total_score += points
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice state before verification
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
