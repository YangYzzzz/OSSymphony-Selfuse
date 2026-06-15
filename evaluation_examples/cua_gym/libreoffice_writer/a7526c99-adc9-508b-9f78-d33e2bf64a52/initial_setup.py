"""
Initial Setup: Configure AutoText meeting minutes template
Task ID: writer_fp_030
Domain: libreoffice_writer

Creates a blank Writer document and opens it in LibreOffice Writer.
No custom AutoText entries exist - the user must create the 'mmtemp' entry.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add some basic content to the document so it's not completely empty
    # This represents a work-in-progress document where the user wants to
    # set up an AutoText shortcut for meeting minutes templates

    heading = doc.add_heading('Project Documentation', level=1)

    doc.add_paragraph(
        'This document serves as the central repository for all project-related '
        'meeting notes and documentation. Team members should use standardized '
        'templates when recording meeting minutes to ensure consistency across '
        'all departments.'
    )

    doc.add_heading('Meeting Schedule', level=2)

    # Add a simple table with upcoming meetings
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Date', 'Meeting Type', 'Organizer']
    for col, h in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    meetings = [
        ['2025-04-07', 'Sprint Planning', 'Sarah Chen'],
        ['2025-04-09', 'Design Review', 'Marcus Johnson'],
        ['2025-04-11', 'Stakeholder Update', 'Priya Patel'],
        ['2025-04-14', 'Retrospective', 'David Kim'],
    ]

    for r, row_data in enumerate(meetings, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacing

    doc.add_heading('Notes', level=2)
    doc.add_paragraph(
        'Please ensure all meeting minutes follow the standard template format. '
        'Contact the project coordinator if you need assistance with document '
        'formatting or template setup.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Ensure LibreOffice profile exists by launching and waiting
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
