"""
Reward Script: Fundraising Event Program - Annual Gala Dinner 2025
Task ID: writer_wf_035
Domain: libreoffice_writer
Scoring:
  C1: Title 22pt bold centered (0.15)
  C2: Date/venue info (0.10)
  C3: Italic welcome message (0.15)
  C4: Schedule table with header + 6 entries (0.20)
  C5: Sponsor categories Gold/Silver/Bronze with 2 each (0.20)
  C6: Thank you note at end (0.10)
  C7: Page border (0.10)
"""

import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_035'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have content (non-empty doc)
    if len(doc.paragraphs) == 0:
        print("FAIL: Document is empty (0 paragraphs)")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title "Annual Gala Dinner 2025" — bold, ~22pt, centered (0.15 pts)
    try:
        found_title = False
        for p in doc.paragraphs:
            if 'annual gala dinner 2025' in p.text.lower():
                # Check centered
                is_centered = (p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
                # Check bold and size across runs
                has_bold_22 = False
                for r in p.runs:
                    if r.font.bold and r.font.size and abs(r.font.size.pt - 22.0) < 1.0:
                        has_bold_22 = True
                        break
                if is_centered and has_bold_22:
                    print(f"PASS: Component 1 — Title is bold, ~22pt, centered (0.15 pts)")
                    total_score += 0.15
                    found_title = True
                else:
                    print(f"FAIL: Component 1 — Title found but centered={is_centered}, bold_22={has_bold_22}")
                    found_title = True
                break
        if not found_title:
            print("FAIL: Component 1 — Title 'Annual Gala Dinner 2025' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Date and venue information present (0.10 pts)
    try:
        has_date_venue = False
        for p in doc.paragraphs:
            text_lower = p.text.lower()
            # Look for date-like info and venue-like info
            if ('2025' in text_lower and ('pm' in text_lower or 'am' in text_lower)) or \
               ('ballroom' in text_lower or 'venue' in text_lower or 'hotel' in text_lower or 'ritz' in text_lower):
                has_date_venue = True
                break
        if has_date_venue:
            print(f"PASS: Component 2 — Date/venue info found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 2 — No date/venue paragraph found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Welcome message — italic paragraph (0.15 pts)
    try:
        has_italic_welcome = False
        for p in doc.paragraphs:
            text_lower = p.text.lower()
            if 'welcome' in text_lower or 'dear' in text_lower or 'pleasure' in text_lower:
                # Check that at least some runs are italic
                italic_runs = [r for r in p.runs if r.font.italic]
                if len(italic_runs) > 0:
                    has_italic_welcome = True
                    print(f"PASS: Component 3 — Italic welcome message found (0.15 pts)")
                    total_score += 0.15
                else:
                    print(f"FAIL: Component 3 — Welcome paragraph found but not italic")
                break
        if not has_italic_welcome and total_score < 0.15 + 0.10 + 0.15:
            # Only print if we haven't already printed pass/fail
            found_any = any('welcome' in p.text.lower() or 'dear' in p.text.lower() for p in doc.paragraphs)
            if not found_any:
                print("FAIL: Component 3 — No welcome message paragraph found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Event schedule table — header + 6 data rows, 2 columns (0.20 pts)
    try:
        has_schedule_table = False
        for table in doc.tables:
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            # Check if it looks like a schedule table (Time + Program Item headers)
            if num_cols >= 2 and num_rows >= 7:
                header_text = ' '.join(c.text.strip().lower() for c in table.rows[0].cells)
                if 'time' in header_text or 'program' in header_text or 'schedule' in header_text:
                    # Count non-empty data rows
                    data_rows = sum(1 for row in table.rows[1:] if any(c.text.strip() for c in row.cells))
                    if data_rows >= 6:
                        print(f"PASS: Component 4 — Schedule table with {num_rows} rows, {num_cols} cols, {data_rows} data entries (0.20 pts)")
                        total_score += 0.20
                        has_schedule_table = True
                    else:
                        print(f"FAIL: Component 4 — Table found but only {data_rows} data rows (need 6)")
                        has_schedule_table = True
                    break
        if not has_schedule_table:
            # Fallback: check any table with >= 7 rows
            for table in doc.tables:
                if len(table.rows) >= 7:
                    print(f"FAIL: Component 4 — Table with {len(table.rows)} rows found but no Time/Program header")
                    has_schedule_table = True
                    break
            if not has_schedule_table:
                print(f"FAIL: Component 4 — No schedule table found (tables: {len(doc.tables)})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Sponsors — Gold, Silver, Bronze with 2 names each (0.20 pts)
    try:
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        full_lower = full_text.lower()

        has_gold = 'gold' in full_lower
        has_silver = 'silver' in full_lower
        has_bronze = 'bronze' in full_lower

        # Count bullet/list items under each category
        # We look for paragraphs with "List Bullet" style or similar after each category heading
        sponsor_count = 0
        current_category = None
        category_counts = {'gold': 0, 'silver': 0, 'bronze': 0}

        for p in doc.paragraphs:
            text_lower = p.text.lower().strip()
            if 'gold' in text_lower and 'sponsor' in text_lower:
                current_category = 'gold'
            elif 'silver' in text_lower and 'sponsor' in text_lower:
                current_category = 'silver'
            elif 'bronze' in text_lower and 'sponsor' in text_lower:
                current_category = 'bronze'
            elif current_category and p.text.strip() and ('List' in (p.style.name or '') or len(p.text.strip()) > 3):
                # This is a sponsor name under the current category
                # Stop counting if we hit another heading-like paragraph
                if p.style and 'List' in p.style.name:
                    category_counts[current_category] += 1
                elif p.text.strip() and not any(kw in text_lower for kw in ['sponsor', 'thank', 'schedule', 'event']):
                    category_counts[current_category] += 1
                else:
                    current_category = None

        total_sponsors = sum(category_counts.values())
        categories_found = sum(1 for v in [has_gold, has_silver, has_bronze] if v)

        score_5 = 0.0
        if categories_found == 3:
            score_5 += 0.10
        elif categories_found >= 2:
            score_5 += 0.05

        if total_sponsors >= 6:
            score_5 += 0.10
        elif total_sponsors >= 4:
            score_5 += 0.05

        if score_5 > 0:
            print(f"PASS: Component 5 — Sponsors: categories={categories_found}/3, names={total_sponsors}/6, per_cat={category_counts} ({score_5} pts)")
            total_score += score_5
        else:
            print(f"FAIL: Component 5 — Sponsors: categories={categories_found}/3, names={total_sponsors}/6")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Thank you note at the end (0.10 pts)
    try:
        # Check last few paragraphs for "thank you" text
        has_thank_you = False
        last_paras = doc.paragraphs[-3:] if len(doc.paragraphs) >= 3 else doc.paragraphs
        for p in last_paras:
            if 'thank' in p.text.lower():
                has_thank_you = True
                break
        if has_thank_you:
            print(f"PASS: Component 6 — Thank you note found at end (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — No thank you note in last paragraphs")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Page border present (0.10 pts)
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        has_border = False
        for section in doc.sections:
            sect_el = section._sectPr
            borders = sect_el.findall('.//w:pgBorders', ns)
            if len(borders) > 0:
                # Verify it has at least top and bottom borders
                for border_el in borders:
                    top = border_el.find('w:top', ns)
                    bottom = border_el.find('w:bottom', ns)
                    if top is not None and bottom is not None:
                        has_border = True
                        break
        if has_border:
            print(f"PASS: Component 7 — Page border found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 — No page border found in section properties")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice edits
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
