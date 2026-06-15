"""
Reward Script: Convert long quotes to proper block quote format
Task ID: writer_acad_036
Domain: libreoffice_writer
Scoring:
  Component 1: All 4 block quote paragraphs have 0.5-inch (457200 EMU) left indent   (0.25 pts)
  Component 2: All 4 block quote paragraphs are single-spaced (line_spacing == 1.0)  (0.25 pts)
  Component 3: All 4 block quote paragraphs have 11pt font size                       (0.25 pts)
  Component 4: All 4 block quote paragraphs have no surrounding quotation marks       (0.25 pts)
Total: 1.0

Ground truth (from context): The four long-quote paragraphs appear at paragraph indices
28, 50, 72, and 94 in the document. In the initial state they are enclosed in double
quotation marks, 12pt, double-spaced (2.0), with zero left indent. The task asks the
agent to remove the quotes and apply: 0.5-inch left indent, single-spacing (1.0), 11pt
font size.

All checks are designed so that they FAIL on initial_env and PASS on golden_env.
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Documents'
TASK_ID = 'writer_acad_036'
FILE_PATH = f'{WORKDIR}/comp_lit_essay.docx'

# Indices of the four block-quote paragraphs in doc.paragraphs (0-based)
BLOCK_QUOTE_INDICES = [28, 50, 72, 94]

# Tolerance for EMU comparison (0.5 inch = 457200 EMU)
EXPECTED_LEFT_INDENT_EMU = 457200
INDENT_TOLERANCE = 1000  # ±1000 EMU (~0.001 inch) tolerance

# Line spacing tolerance
EXPECTED_LINE_SPACING = 1.0
SPACING_TOLERANCE = 0.05

# Font size tolerance
EXPECTED_FONT_PT = 11.0
FONT_TOLERANCE = 0.5  # ±0.5pt

# Quotation mark characters to detect
QUOTE_CHARS = ('"', '\u201c', '\u201d', '\u2018', '\u2019')


def verify_task(file_path):
    """
    Verify that all four long-quote paragraphs have been converted to block quote format.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: file must exist and be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Verify we have enough paragraphs
    if len(doc.paragraphs) <= max(BLOCK_QUOTE_INDICES):
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, "
              f"expected at least {max(BLOCK_QUOTE_INDICES) + 1}")
        print("REWARD: 0.0")
        return 0.0

    # Collect the four block-quote paragraphs
    bq_paras = [doc.paragraphs[i] for i in BLOCK_QUOTE_INDICES]

    # -------------------------------------------------------------------------
    # Component 1: All 4 block quotes have 0.5-inch (457200 EMU) left indent  (0.25 pts)
    # -------------------------------------------------------------------------
    try:
        indent_pass = []
        for i, (idx, para) in enumerate(zip(BLOCK_QUOTE_INDICES, bq_paras)):
            pf = para.paragraph_format
            indent = pf.left_indent
            # indent may be None (no explicit indent) or a numeric EMU value
            if indent is None:
                indent_val = 0
            else:
                indent_val = int(indent)
            ok = abs(indent_val - EXPECTED_LEFT_INDENT_EMU) <= INDENT_TOLERANCE
            indent_pass.append(ok)
            if ok:
                print(f"  PASS (indent): Para {idx} left_indent={indent_val} EMU (~{indent_val/914400:.3f} in)")
            else:
                print(f"  FAIL (indent): Para {idx} left_indent={indent_val} EMU, "
                      f"expected ~{EXPECTED_LEFT_INDENT_EMU} EMU (0.5 in)")

        if all(indent_pass):
            print(f"PASS: Component 1 — All 4 block quotes have 0.5-inch left indent (0.25 pts)")
            total_score += 0.25
        else:
            fails = sum(1 for x in indent_pass if not x)
            print(f"FAIL: Component 1 — {fails}/4 block quote(s) missing 0.5-inch left indent")
    except Exception as e:
        print(f"ERROR: Component 1 (left indent check) — {e}")

    # -------------------------------------------------------------------------
    # Component 2: All 4 block quotes are single-spaced (line_spacing == 1.0)  (0.25 pts)
    # -------------------------------------------------------------------------
    try:
        spacing_pass = []
        for i, (idx, para) in enumerate(zip(BLOCK_QUOTE_INDICES, bq_paras)):
            pf = para.paragraph_format
            spacing = pf.line_spacing
            if spacing is None:
                spacing_val = None
            elif hasattr(spacing, 'pt'):
                # If it's an Pt value, convert to multiplier (baseline 12pt body = 12pt)
                spacing_val = float(spacing.pt) / 12.0
            else:
                spacing_val = float(spacing)
            ok = (spacing_val is not None and
                  abs(spacing_val - EXPECTED_LINE_SPACING) <= SPACING_TOLERANCE)
            spacing_pass.append(ok)
            if ok:
                print(f"  PASS (spacing): Para {idx} line_spacing={spacing_val}")
            else:
                print(f"  FAIL (spacing): Para {idx} line_spacing={spacing_val}, expected {EXPECTED_LINE_SPACING}")

        if all(spacing_pass):
            print(f"PASS: Component 2 — All 4 block quotes are single-spaced (0.25 pts)")
            total_score += 0.25
        else:
            fails = sum(1 for x in spacing_pass if not x)
            print(f"FAIL: Component 2 — {fails}/4 block quote(s) not single-spaced")
    except Exception as e:
        print(f"ERROR: Component 2 (line spacing check) — {e}")

    # -------------------------------------------------------------------------
    # Component 3: All 4 block quotes have 11pt font size                      (0.25 pts)
    # -------------------------------------------------------------------------
    try:
        font_pass = []
        for i, (idx, para) in enumerate(zip(BLOCK_QUOTE_INDICES, bq_paras)):
            para_font_sizes = []
            for run in para.runs:
                if run.font.size is not None:
                    para_font_sizes.append(run.font.size.pt)
            if not para_font_sizes:
                # No explicit run-level font size; treat as not meeting requirement
                ok = False
                print(f"  FAIL (font): Para {idx} — no explicit run font sizes found")
            else:
                # All runs with explicit size should be ~11pt
                ok = all(abs(sz - EXPECTED_FONT_PT) <= FONT_TOLERANCE for sz in para_font_sizes)
                avg = sum(para_font_sizes) / len(para_font_sizes)
                if ok:
                    print(f"  PASS (font): Para {idx} — run font sizes: {para_font_sizes} pt (avg={avg:.1f})")
                else:
                    bad = [sz for sz in para_font_sizes if abs(sz - EXPECTED_FONT_PT) > FONT_TOLERANCE]
                    print(f"  FAIL (font): Para {idx} — unexpected sizes: {bad}, expected ~{EXPECTED_FONT_PT}pt")
            font_pass.append(ok)

        if all(font_pass):
            print(f"PASS: Component 3 — All 4 block quotes have 11pt font size (0.25 pts)")
            total_score += 0.25
        else:
            fails = sum(1 for x in font_pass if not x)
            print(f"FAIL: Component 3 — {fails}/4 block quote(s) not at 11pt font size")
    except Exception as e:
        print(f"ERROR: Component 3 (font size check) — {e}")

    # -------------------------------------------------------------------------
    # Component 4: All 4 block quotes have no leading/trailing quotation marks (0.25 pts)
    # -------------------------------------------------------------------------
    try:
        noquote_pass = []
        for i, (idx, para) in enumerate(zip(BLOCK_QUOTE_INDICES, bq_paras)):
            text = para.text.strip()
            if not text:
                noquote_pass.append(False)
                print(f"  FAIL (no-quote): Para {idx} — empty text")
                continue
            starts_with_quote = text[0] in QUOTE_CHARS
            ends_with_quote = text[-1] in QUOTE_CHARS
            ok = not starts_with_quote and not ends_with_quote
            noquote_pass.append(ok)
            if ok:
                print(f"  PASS (no-quote): Para {idx} — no surrounding quote chars "
                      f"(starts: '{text[0]}', ends: '{text[-1]}')")
            else:
                print(f"  FAIL (no-quote): Para {idx} — quote chars still present "
                      f"(starts: '{text[0]}', ends: '{text[-1]}')")

        if all(noquote_pass):
            print(f"PASS: Component 4 — All 4 block quotes have quotation marks removed (0.25 pts)")
            total_score += 0.25
        else:
            fails = sum(1 for x in noquote_pass if not x)
            print(f"FAIL: Component 4 — {fails}/4 block quote(s) still have surrounding quotation marks")
    except Exception as e:
        print(f"ERROR: Component 4 (no quotation marks check) — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
