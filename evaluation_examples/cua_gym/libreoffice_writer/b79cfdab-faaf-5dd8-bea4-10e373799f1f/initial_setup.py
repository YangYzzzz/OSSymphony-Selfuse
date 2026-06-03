"""
Initial Setup: Comparative Literature Essay with 4 long quotes in regular paragraph format
Task ID: writer_acad_036
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Documents'
TASK_ID = 'writer_acad_036'
OUTPUT = f'{WORKDIR}/comp_lit_essay.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_para(doc, text):
    """Add a standard body paragraph: 12pt, double-spaced, no indent."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    pf = para.paragraph_format
    pf.line_spacing = 2.0
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return para


def add_heading_para(doc, text, level=1):
    """Add a section heading."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    pf = para.paragraph_format
    pf.line_spacing = 2.0
    pf.left_indent = Inches(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return para


def add_quote_para_initial(doc, text):
    """
    Add a long quote in INITIAL state: regular paragraph format,
    with double quotation marks, 12pt, double-spaced, no indent.
    This is what the agent needs to convert to block quote format.
    """
    para = doc.add_paragraph()
    # Add the text with surrounding quotation marks
    run = para.add_run(f'"{text}"')
    run.font.size = Pt(12)
    pf = para.paragraph_format
    pf.line_spacing = 2.0
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return para


def create_initial():
    # Ensure Documents directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set default section margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- TITLE PAGE (page 1) ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(
        "Thresholds of Language: Alienation and Belonging in\n"
        "Conrad, Kafka, and Beckett"
    )
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.line_spacing = 2.0
    title_para.paragraph_format.space_before = Pt(72)

    for _ in range(8):
        add_body_para(doc, "")

    author_para = doc.add_paragraph()
    author_run = author_para.add_run(
        "Emily R. Thornton\n"
        "ENGL 4850: Comparative Modernisms\n"
        "Professor Adriana Vasquez\n"
        "March 14, 2025"
    )
    author_run.font.size = Pt(12)
    author_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.line_spacing = 2.0

    doc.add_page_break()

    # --- PAGE 2: Introduction ---
    add_heading_para(doc, "Introduction")
    add_body_para(doc, "")

    add_body_para(doc,
        "The question of how language simultaneously enables and forecloses identity has "
        "preoccupied literary modernism since its inception. For writers navigating the contested "
        "terrain between cultures, languages, and historical moments, the act of writing in a "
        "second or adopted language becomes not merely a stylistic choice but an ontological "
        "condition. Joseph Conrad, Franz Kafka, and Samuel Beckett each confronted this condition "
        "in radically different yet curiously convergent ways, producing a body of work that "
        "interrogates the very foundations of literary expression."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "This essay argues that the experience of linguistic estrangement in Conrad, Kafka, and "
        "Beckett does not merely produce a literature of loss but generates a productive tension "
        "between belonging and alienation that becomes the defining formal and thematic feature of "
        "their work. By examining key passages from Heart of Darkness, The Trial, and Waiting for "
        "Godot, alongside secondary scholarship by critics including Edward Said, Marjorie Perloff, "
        "and Michael Wood, this analysis will demonstrate how each author transforms the condition "
        "of linguistic exile into a mode of literary inquiry."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The three authors under consideration occupy distinct positions within the modernist "
        "tradition, yet share a formative experience of writing at the margins of dominant literary "
        "culture. Conrad, born Józef Teodor Konrad Korzeniowski in the Polish Ukraine, learned "
        "English as his third language, producing some of the most stylistically distinctive prose "
        "in the English canon. Kafka wrote in German while living in Prague, a city where language "
        "itself was a site of political and ethnic contestation. Beckett, an Irishman who chose to "
        "write primarily in French, described his adoption of a second language as a way of writing "
        "without style, stripped of the ornamental confidence of a native tongue."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Previous scholarship has tended to treat these three figures separately, addressing "
        "Conrad's imperialism, Kafka's bureaucratic surrealism, and Beckett's existential minimalism "
        "as distinct phenomena. Yet read together, through the lens of linguistic alienation, their "
        "works reveal a shared poetics of estrangement that has had profound implications for "
        "subsequent literary production. The analysis that follows seeks to illuminate this shared "
        "poetics while attending carefully to the historical and biographical specificities that "
        "differentiate each author's relationship to language."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 3: Section I — Conrad ---
    add_heading_para(doc, "I. Joseph Conrad and the Burden of the Adopted Tongue")
    add_body_para(doc, "")

    add_body_para(doc,
        "Conrad's position within English literary history has long been defined by a paradox: "
        "the writer who produced some of the most celebrated English prose was not a native "
        "speaker of the language. Henry James famously remarked that Conrad's English read as if "
        "translated from some other, stranger tongue. This observation, intended as a mild "
        "reservation, has since become a touchstone for discussions of Conrad's style and his "
        "relationship to the literary tradition he both inhabited and transformed."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "For Conrad, the decision to write in English rather than French, which he also spoke "
        "fluently, was not simply practical. As he explained in A Personal Record, English "
        "represented for him a kind of neutral ground, free from the heavy freight of national "
        "memory that both Polish and French carried. Yet this neutrality was itself a fiction, "
        "and Conrad's prose repeatedly stages the inadequacy of any language to fully capture "
        "the experience it seeks to represent. Marlow's famous observation in Heart of Darkness "
        "about the difficulty of conveying experience to another person stands as a self-reflexive "
        "comment on the impossibility of Conrad's own literary enterprise."
    )
    add_body_para(doc, "")

    # QUOTE 1 (page 3): ~55 words
    add_quote_para_initial(doc,
        "Do you see the story? Do you see anything? It seems to me I am trying to tell you a "
        "dream — making a vain attempt, because no relation of a dream can convey the "
        "dream-sensation, that commingling of absurdity, surprise, and bewilderment in a tremor "
        "of struggling revolt, that notion of being captured by the incredible which is of the "
        "very essence of dreams."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "This passage, in which Marlow explicitly acknowledges the limits of narrative "
        "transmission, encapsulates Conrad's broader preoccupation with the gap between "
        "experience and representation. The dream metaphor is significant: dreams are private, "
        "untranslatable, resistant to the public medium of language. Yet Marlow, and by "
        "extension Conrad, persists in the attempt, suggesting that the value of the literary "
        "enterprise lies precisely in the effort rather than the achievement. Edward Said, in "
        "his influential reading of Conrad, argues that this formal self-consciousness reflects "
        "the condition of the exile who can never fully inhabit the culture whose language "
        "he has adopted."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The formal consequence of this condition is what critics have identified as Conrad's "
        "signature indirection: the use of multiple narrators, embedded narratives, and "
        "deliberate obscurity that forces the reader to participate actively in the construction "
        "of meaning. This technique, which Marlow himself theorizes in the opening pages of "
        "Heart of Darkness, is not a stylistic tic but a structural response to the impossibility "
        "of transparent communication across the boundaries of language and experience."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 4: Conrad continued ---
    add_body_para(doc,
        "Conrad's relationship to English was further complicated by his awareness of the "
        "political dimensions of language. Writing at the height of British imperial power, "
        "Conrad was acutely conscious that English was not merely a medium of literary expression "
        "but an instrument of colonial authority. The language in which he chose to write was "
        "also the language in which the empire conducted its business, issued its orders, and "
        "justified its existence. This awareness saturates Heart of Darkness, where the "
        "resonant abstractions of imperial rhetoric — civilization, progress, duty — are "
        "subjected to a withering ironic scrutiny."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Critics have debated whether Conrad's irony constitutes a genuine critique of "
        "imperialism or merely a aesthetic distancing that ultimately reproduces the imperial "
        "gaze it claims to interrogate. Chinua Achebe's famous attack on Heart of Darkness as "
        "a racist text opened a sustained critical controversy that has never been fully resolved. "
        "Yet even Achebe acknowledged the power of Conrad's prose, and it is perhaps in the "
        "tension between Conrad's critical intelligence and his complicity in imperial discourse "
        "that we find the most productive site of analysis. The linguistic exile who writes in "
        "the language of empire cannot escape the contradictions that condition inhabits."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "More recently, postcolonial scholars have revisited Conrad through the lens of "
        "translation theory, arguing that his prose enacts a kind of perpetual translation that "
        "unsettles the comfortable assumptions of monolingual literary culture. Frances Singh "
        "and later Benita Parry have suggested that Conrad's stylistic opacity is a form of "
        "resistance, however ambivalent, to the transparency that imperial communication "
        "demands. This reading has found support in biographical scholarship that emphasizes "
        "Conrad's own experience of linguistic and cultural displacement as the enabling "
        "condition of his literary achievement."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Whatever one makes of the political dimensions of Conrad's work, it is clear that "
        "his relationship to English shaped not only the content but the form of his fiction in "
        "profound and distinctive ways. The tortured syntax, the layered narrative perspectives, "
        "the obsessive return to moments of communicative failure — all of these features bear "
        "the marks of a consciousness formed in one language and forced to express itself in "
        "another. Conrad's English is not simply English with a foreign accent; it is a "
        "language transformed by the pressure of everything that cannot be said in it."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 5: Section II — Kafka ---
    add_heading_para(doc, "II. Franz Kafka and the Politics of Linguistic Marginality")
    add_body_para(doc, "")

    add_body_para(doc,
        "Kafka's relationship to language was shaped by a different set of historical "
        "pressures. Writing in Prague at the turn of the twentieth century, Kafka inhabited "
        "a city where language was inseparable from ethnic identity and political allegiance. "
        "The German of Prague's Jewish intellectual community was already a kind of minority "
        "language, distinguished from the German of Germany by subtle but significant "
        "differences of usage and association. Kafka's awareness of this marginality inflects "
        "his prose in ways that have only recently begun to receive adequate critical attention."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Gilles Deleuze and Félix Guattari, in their influential study of Kafka, introduced "
        "the concept of minor literature to describe the condition of writing in a major "
        "language from a position of cultural marginality. Kafka's German, they argued, was "
        "minor not in the sense of inferior but in the sense of being inhabited differently "
        "by someone for whom it was never simply home. This notion of minor literature has "
        "been enormously productive for subsequent scholarship, though it has also been "
        "criticized for romanticizing a condition that Kafka himself experienced with "
        "considerable ambivalence."
    )
    add_body_para(doc, "")

    # QUOTE 2 (page 5): ~60 words
    add_quote_para_initial(doc,
        "The German language is the medium through which I express myself, but it is not my "
        "own; I can handle it with some facility, but it does not yield the innermost resources "
        "of the self. One writes from need, from compulsion, from a kind of obscure necessity "
        "whose origins are not fully apparent to the writer himself, and the language that "
        "serves this need is always, in some fundamental sense, borrowed, provisional, not "
        "entirely one's own."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "This observation — drawn from Kafka's diaries and letters rather than his fiction — "
        "suggests that Kafka's experience of linguistic marginality was not merely a "
        "sociological condition but an existential one. The language in which he wrote was "
        "also the language in which he thought, dreamed, and conducted his inner life, yet "
        "it remained somehow external to him, a medium he had mastered without ever fully "
        "claiming as his own. This peculiar relationship to language generates much of the "
        "unsettling quality of Kafka's prose, its sense of precision applied to situations "
        "that resist precise definition."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The Trial offers perhaps the most sustained exploration of this theme. Josef K., "
        "the protagonist, finds himself subject to a legal process conducted entirely in a "
        "language he does not understand, governed by rules he has no access to, adjudicated "
        "by officials who speak with authority about matters they refuse to clarify. The "
        "allegorical dimensions of this scenario are clear: Kafka is mapping the condition of "
        "the linguistic outsider onto the condition of the legal subject, suggesting that "
        "the experience of not quite belonging to the language of authority is a form of "
        "permanent legal minority. One is always already accused by a law one cannot read."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 6: Kafka continued ---
    add_body_para(doc,
        "Kafka's formal innovations are inseparable from this thematic preoccupation. "
        "The flat, affectless prose style that characterizes The Trial and The Metamorphosis "
        "has often been described as a kind of bureaucratic German, the language of official "
        "reports and administrative memoranda applied to situations of extreme existential "
        "crisis. This juxtaposition of flat style and extreme content is itself a form of "
        "estrangement, forcing the reader into an uncomfortable confrontation with the "
        "inadequacy of conventional language to represent unconventional experience."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The famous opening sentence of The Metamorphosis — in Willa and Edwin Muir's "
        "translation, 'As Gregor Samsa awoke one morning from uneasy dreams he found himself "
        "transformed in his bed into a gigantic insect' — is remarkable precisely for its "
        "refusal of the extraordinary. The transformation is presented in the same syntactic "
        "structure as any other waking experience, as if the grammar of ordinary life were "
        "adequate to describe the most radical discontinuity of being. This grammatical "
        "indifference to the extraordinary is Kafka's signature gesture, and it is deeply "
        "connected to his experience of linguistic marginality."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Marjorie Perloff has argued that Kafka's prose style represents a kind of "
        "willed averageness, a systematic refusal of the rhetorical ornament that would "
        "mark it as the product of a confident native speaker. This reading suggests that "
        "Kafka's minimalism is not simply an aesthetic choice but a response to his "
        "linguistic situation, a way of writing that acknowledges its own provisional "
        "relationship to the language it inhabits. The result is a prose style of remarkable "
        "power precisely because of its apparent impoverishment, its refusal of the richness "
        "that the language could, in other hands, provide."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The political dimensions of Kafka's linguistic position have become increasingly "
        "prominent in recent scholarship. The Prague that Kafka inhabited was a city in "
        "political transformation, moving toward the dissolution of the Habsburg Empire and "
        "the emergence of new national formations. In this context, the question of which "
        "language one wrote in was not merely a literary choice but a political declaration. "
        "Kafka's decision to write in German placed him on one side of a complex series "
        "of ethnic, linguistic, and political divisions, even as his prose seemed to "
        "withdraw from all declarations into a space of studied neutrality."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 7: Section III — Beckett ---
    add_heading_para(doc, "III. Samuel Beckett and the Aesthetics of Linguistic Dispossession")
    add_body_para(doc, "")

    add_body_para(doc,
        "Of the three writers under consideration, Beckett presents the most radical case "
        "of deliberate linguistic displacement. Unlike Conrad and Kafka, who wrote in adopted "
        "or minority languages out of biographical necessity, Beckett made a conscious choice "
        "to abandon his native English in favor of French. This choice, made in the late "
        "1940s after years of writing in English, represents one of the most consequential "
        "and theorized acts of linguistic self-displacement in twentieth-century literary "
        "history."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Beckett described his turn to French in various ways at various times, but the "
        "most frequently cited explanation is that writing in French allowed him to write "
        "without style, to strip away the verbal exuberance of his early English prose and "
        "arrive at something more essential, more stripped of the ornamental confidence "
        "that a native speaker brings to his language. This explanation has been enormously "
        "influential in shaping the reception of Beckett's work, though it has also been "
        "subjected to considerable critical scrutiny."
    )
    add_body_para(doc, "")

    # QUOTE 3 (page 7): ~68 words
    add_quote_para_initial(doc,
        "I chose French because you can't do tricks with it, can't manipulate it as you can "
        "English, use it as a conjurer uses his props. I needed the bareness, the poverty "
        "of the means. I needed to find a language without resonance, without those "
        "innumerable overtones that hang about every English word like a private history "
        "you haven't been invited to share. French gave me distance from my subject, and "
        "distance was what I required to see it clearly, to stop performing and simply "
        "witness what remained when the performance was stripped away."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "This account of his linguistic choice reveals Beckett's preoccupation with a "
        "problem that is at once aesthetic and philosophical: the problem of how to use "
        "language to gesture toward what language cannot contain. The resonances, "
        "overtones, and private histories that he attributes to English are not accidental "
        "features of the language but constitutive ones; they are what make English "
        "English, what give it its expressive resources. By choosing a language that "
        "lacked these resources for him, Beckett was not impoverishing his writing but "
        "redirecting it toward a different kind of expressiveness — one based on what "
        "is withheld rather than what is said."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Waiting for Godot, written in French in 1948-1949 and subsequently translated "
        "into English by Beckett himself, represents the fullest realization of this "
        "aesthetic. The play's dialogue is structured around repetition, non-sequitur, "
        "and the failure of communication, with Vladimir and Estragon conducting a "
        "conversation that is simultaneously about nothing and about the fundamental "
        "conditions of human existence. The minimalism of the language — short sentences, "
        "simple vocabulary, frequent pauses and silences — is not a poverty of means but "
        "an achievement of radical compression."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 8: Beckett continued ---
    add_body_para(doc,
        "The translation of Godot from French to English raises the question of what "
        "happens when the language of dispossession is translated back into the language "
        "of origin. Beckett's English version is not simply a translation but a new work, "
        "altered by the process of moving between languages. The differences between the "
        "two versions have been carefully documented by scholars, and they reveal the "
        "extent to which meaning is not simply transferred between languages but transformed "
        "in the act of transfer. Beckett's self-translations are among the most fascinating "
        "documents of twentieth-century literary linguistics."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The formal structure of Godot enacts the impossibility of completion that "
        "characterizes Beckett's broader aesthetic. The play ends as it began, with "
        "Vladimir and Estragon waiting, unable to leave, unable to act, sustained by "
        "the possibility of an arrival that never occurs. This circular structure is "
        "formally homologous to the condition of the linguistic exile: one is perpetually "
        "between languages, neither fully at home in the adopted tongue nor able to return "
        "to the native one. The waiting is not passive but constitutive; it is the "
        "condition from which Beckett's art proceeds."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Michael Wood, in his study of literature and the taste of words, has argued that "
        "Beckett's linguistic minimalism represents a form of what he calls the "
        "deliberate cultivation of speechlessness, a systematic reduction of language "
        "to its irreducible minimum in order to reveal what lies on the other side of "
        "expression. This reading suggests that Beckett's turn to French was not simply "
        "a biographical accident but a principled aesthetic strategy, a way of using "
        "the resources of linguistic displacement to arrive at a form of utterance that "
        "transcends the limitations of any particular language."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The critical literature on Beckett's bilingualism is by now substantial, and "
        "its conclusions are broadly convergent: writing between languages produces a "
        "literature of the gap, a literature that inhabits the space between what can "
        "be said in one language and what can be said in another. This space is not "
        "empty but richly populated by the ghosts of meaning that can never be fully "
        "domesticated in either tongue. Beckett's great achievement was to find formal "
        "structures adequate to this condition, structures that make the impossibility "
        "of full expression itself the subject and substance of literary art."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 9: Section IV — Comparative Analysis ---
    add_heading_para(doc, "IV. Shared Poetics: Estrangement, Form, and the Limits of Language")
    add_body_para(doc, "")

    add_body_para(doc,
        "Having examined the three writers individually, we are now in a position to "
        "identify the structural commonalities that link their very different literary "
        "projects. What Conrad, Kafka, and Beckett share is not simply a biographical "
        "condition of linguistic displacement but a formal response to that condition: "
        "a set of stylistic strategies that convert the experience of not-quite-belonging "
        "to language into literary form. These strategies differ in their particulars, "
        "but they share a common orientation toward the limits of language as a productive "
        "literary resource."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The most fundamental of these strategies is what we might call productive "
        "indirection: the use of narrative or formal structures that approach meaning "
        "obliquely rather than directly. In Conrad, this takes the form of layered "
        "narration and deliberate obscurity. In Kafka, it manifests as bureaucratic "
        "literalism applied to metaphysical situations. In Beckett, it appears as radical "
        "compression and the foregrounding of what cannot be said. All three strategies "
        "exploit the gap between language and experience that the condition of linguistic "
        "displacement makes visible."
    )
    add_body_para(doc, "")

    # QUOTE 4 (page 9): ~72 words
    add_quote_para_initial(doc,
        "The problem of language for the literary exile is not simply that words are "
        "inadequate to experience, though they are. It is that the very inadequacy "
        "of language becomes itself a form of experience, a condition that shapes "
        "consciousness and demands formal response. When Conrad reaches for the word "
        "that will convey precisely what Marlow cannot convey, when Kafka deploys "
        "administrative precision to describe the incomprehensible, when Beckett reduces "
        "language to its bare bones and then reduces it further still, they are all "
        "performing the same fundamental literary act: making the failure of language "
        "into a vehicle of meaning."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "This shared poetics of estrangement has had lasting consequences for literary "
        "history. The techniques that Conrad, Kafka, and Beckett developed in response "
        "to their particular conditions of linguistic displacement became available to "
        "subsequent writers as formal resources, detached from the specific biographical "
        "circumstances that generated them. Writers as different as Toni Morrison, "
        "Salman Rushdie, and W. G. Sebald have drawn on the tradition these three "
        "writers established, finding in the aesthetics of linguistic exile a set of "
        "tools for articulating their own diverse experiences of cultural displacement."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The theoretical implications of this tradition are equally significant. "
        "The recognition that linguistic marginality can generate distinctive literary "
        "forms challenges the assumption, still common in some quarters of literary "
        "criticism, that literature is best understood as the expression of a unified "
        "cultural tradition. Conrad, Kafka, and Beckett wrote from positions of "
        "cultural multiplicity, and their work reflects and enacts the productive "
        "tensions of that multiplicity. To read their work without attending to these "
        "tensions is to miss much of what makes it powerful and enduring."
    )
    add_body_para(doc, "")

    doc.add_page_break()

    # --- PAGE 10: Conclusion ---
    add_heading_para(doc, "Conclusion")
    add_body_para(doc, "")

    add_body_para(doc,
        "This essay has argued that Joseph Conrad, Franz Kafka, and Samuel Beckett, "
        "despite their very different biographical and historical situations, share a "
        "poetics of linguistic estrangement that shapes both the form and content of "
        "their major works. By examining key passages from Heart of Darkness, The Trial, "
        "and Waiting for Godot alongside relevant secondary scholarship, the analysis "
        "has demonstrated how each author transforms the condition of writing in an "
        "adopted or minority language from a biographical limitation into a source of "
        "distinctive literary power."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The implications of this analysis extend beyond the three figures examined here. "
        "The poetics of linguistic estrangement that Conrad, Kafka, and Beckett developed "
        "has become one of the defining features of modernist and postmodernist literature "
        "more broadly, shaping the work of writers across languages, cultures, and "
        "historical periods. To understand this poetics in its original formulations is "
        "to gain insight not only into the work of these three major authors but into "
        "the conditions and possibilities of literary expression in a world of multiple "
        "languages and cultures."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "Further research in this area might usefully examine how the specific linguistic "
        "and cultural contexts of Conrad's, Kafka's, and Beckett's writing have been "
        "received and transformed in translation and in the various national literary "
        "traditions that have claimed them. The question of how literary estrangement "
        "travels — how the poetics of the linguistic outsider is received by different "
        "reading communities with different relationships to the languages involved — "
        "remains, to this point, underexplored. It is a question that has implications "
        "not only for literary history but for our understanding of how literature "
        "functions as a medium of cultural exchange."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "What Conrad, Kafka, and Beckett ultimately share is a conviction — borne out "
        "in the formal achievements of their major works — that the limits of language "
        "are not simply obstacles to expression but the very site at which literary art "
        "becomes possible. The gap between experience and representation, between what "
        "one knows and what one can say, is not a failure of language but its most "
        "productive condition. It is from within this gap that their writing emerges, "
        "and it is to this gap that their work continues to return, finding in it an "
        "inexhaustible resource for literary inquiry."
    )
    add_body_para(doc, "")

    add_body_para(doc,
        "The comparative framework employed in this essay has certain limitations that "
        "should be acknowledged. The focus on linguistic displacement as the primary "
        "analytic lens necessarily foregrounds certain aspects of these authors' work "
        "while backgrounding others. A fuller account would need to attend more "
        "carefully to the specific historical contexts — imperial Britain, Habsburg "
        "Prague, postwar Paris — that shaped each author's experience of linguistic "
        "marginality. It would also need to engage more directly with the debates within "
        "postcolonial and minority literature scholarship about the political dimensions "
        "of writing in borrowed tongues."
    )
    add_body_para(doc, "")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
