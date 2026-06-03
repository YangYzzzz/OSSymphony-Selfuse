"""
Reward Script: Change document orientation from portrait to landscape
Task ID: writer_hr_020
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Orientation enum is LANDSCAPE
  Component 2 (0.3): Page dimensions are swapped (width > height)
  Component 3 (0.2): Orientation is LANDSCAPE AND document content preserved
"""

import os
import time


WORKDIR = '/home/user'
TASK_ID = 'writer_hr_020'


def persist_app_state(domain: str):
    """Save any unsaved GUI state before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Emu
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # We only check section 0 (the document has one section)
    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: Orientation enum is LANDSCAPE (0.5 points)
    try:
        orientation = section.orientation
        if orientation == WD_ORIENT.LANDSCAPE:
            print(f"PASS: Component 1 — Orientation is LANDSCAPE ({orientation}) (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Expected LANDSCAPE, found {orientation}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Page dimensions are swapped for landscape (width > height) (0.3 points)
    # Portrait standard: width=8.5in, height=11in
    # Landscape standard: width=11in, height=8.5in
    try:
        page_width = section.page_width
        page_height = section.page_height
        width_inches = Emu(page_width).inches
        height_inches = Emu(page_height).inches

        if page_width > page_height:
            print(f"PASS: Component 2 — Width ({width_inches:.2f}in) > Height ({height_inches:.2f}in), landscape dimensions (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Width ({width_inches:.2f}in) <= Height ({height_inches:.2f}in), not landscape dimensions")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Orientation is LANDSCAPE AND document content preserved (0.2 points)
    # This compound check ensures the orientation change didn't corrupt the document.
    # Gates on landscape so it won't score on the initial (portrait) document.
    try:
        is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
        has_title = any(p.text.strip() == "Org_Chart_2026" for p in doc.paragraphs)
        has_tables = len(doc.tables) >= 3
        has_content = len(doc.paragraphs) >= 5

        if is_landscape and has_title and has_tables and has_content:
            print(f"PASS: Component 3 — Landscape AND content preserved (title='Org_Chart_2026', tables={len(doc.tables)}, paragraphs={len(doc.paragraphs)}) (0.2 pts)")
            total_score += 0.2
        else:
            details = []
            if not is_landscape:
                details.append("not landscape")
            if not has_title:
                details.append("title 'Org_Chart_2026' missing")
            if not has_tables:
                details.append(f"only {len(doc.tables)} tables (need >=3)")
            if not has_content:
                details.append(f"only {len(doc.paragraphs)} paragraphs (need >=5)")
            print(f"FAIL: Component 3 — {', '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'

persist_app_state("libreoffice_writer")

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
