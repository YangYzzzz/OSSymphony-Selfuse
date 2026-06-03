"""
Initial Setup: Create organizational chart document in portrait orientation
Task ID: writer_hr_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_020'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup: Portrait (default) ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('Org_Chart_2026', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Pinnacle Solutions Inc. - Organizational Structure')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.italic = True

    doc.add_paragraph()  # spacer

    # --- CEO Level ---
    p_ceo = doc.add_paragraph()
    p_ceo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_ceo = p_ceo.add_run('CEO: Margaret Thornton')
    run_ceo.bold = True
    run_ceo.font.size = Pt(16)
    run_ceo.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()  # spacer

    # --- Organizational Chart Table (wide, to motivate landscape) ---
    # 6 columns for 6 departments - intentionally wide
    departments = [
        'Engineering\nVP: David Kim',
        'Marketing\nVP: Lisa Chen',
        'Finance\nVP: Robert Hayes',
        'Operations\nVP: Priya Patel',
        'Human Resources\nVP: James Wilson',
        'Legal & Compliance\nVP: Angela Torres',
    ]

    # VP-level row
    table_vp = doc.add_table(rows=1, cols=6)
    table_vp.style = 'Table Grid'
    for i, dept in enumerate(departments):
        cell = table_vp.cell(0, i)
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(dept)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()  # spacer

    # --- Director Level Table ---
    directors = [
        ['Dir: Sarah Nguyen\n(Software Dev)',
         'Dir: Mark Thompson\n(Brand Strategy)',
         'Dir: Emily Rodriguez\n(Accounting)',
         'Dir: Anil Sharma\n(Supply Chain)',
         'Dir: Rachel Foster\n(Talent Acquisition)',
         'Dir: Thomas Wright\n(Corporate Law)'],
        ['Dir: Michael Park\n(Infrastructure)',
         'Dir: Olivia Martinez\n(Digital Marketing)',
         'Dir: Brian Cooper\n(FP&A)',
         'Dir: Fatima Al-Hassan\n(Logistics)',
         'Dir: Daniel Kim\n(Employee Relations)',
         'Dir: Catherine Lee\n(Regulatory Affairs)'],
    ]

    table_dir = doc.add_table(rows=2, cols=6)
    table_dir.style = 'Table Grid'
    for r, row_data in enumerate(directors):
        for c, text in enumerate(row_data):
            cell = table_dir.cell(r, c)
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.add_run(text)
            run.font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # --- Manager/Team Lead Level Table ---
    managers = [
        ['Team Lead:\nAlex Rivera\n(Backend)',
         'Team Lead:\nJessica Brown\n(Content)',
         'Sr. Analyst:\nKevin Zhang\n(Budget)',
         'Manager:\nMaria Santos\n(Warehousing)',
         'Recruiter:\nNatalie Green\n(Tech Hiring)',
         'Paralegal:\nSteven Moore\n(Contracts)'],
        ['Team Lead:\nChris Taylor\n(Frontend)',
         'Team Lead:\nAmanda White\n(Social Media)',
         'Sr. Analyst:\nJulia Watson\n(Audit)',
         'Manager:\nOmar Hassan\n(Procurement)',
         'Specialist:\nLaura Adams\n(Benefits)',
         'Counsel:\nPatrick Quinn\n(IP Law)'],
        ['Team Lead:\nRyan Lee\n(DevOps)',
         'Designer:\nSophia Garcia\n(Creative)',
         'Controller:\nAndrew Davis\n(Tax)',
         'Coordinator:\nEric Johnson\n(Shipping)',
         'Analyst:\nMegan Clark\n(Training)',
         'Analyst:\nDiana Miller\n(Compliance)'],
    ]

    table_mgr = doc.add_table(rows=3, cols=6)
    table_mgr.style = 'Table Grid'
    for r, row_data in enumerate(managers):
        for c, text in enumerate(row_data):
            cell = table_mgr.cell(r, c)
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.add_run(text)
            run.font.size = Pt(8)

    doc.add_paragraph()  # spacer

    # --- Footer note ---
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_note = footer_para.add_run(
        'Note: This organizational chart reflects the structure as of January 2026. '
        'For updates, contact the HR department at hr@pinnaclesolutions.com. '
        'Dotted-line reporting relationships are not shown in this version.'
    )
    run_note.font.size = Pt(9)
    run_note.italic = True
    run_note.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # --- Effective date ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_date = date_para.add_run('Effective: January 15, 2026')
    run_date.font.size = Pt(9)
    run_date.bold = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
