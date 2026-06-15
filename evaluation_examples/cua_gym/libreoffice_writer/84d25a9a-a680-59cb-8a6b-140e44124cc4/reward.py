"""
Reward Script: Create a 3x3 tic-tac-toe table in a Word document
Task ID: writer_tbl_015
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): 3x3 table exists with all empty cells
  Component 2 (0.30): Each cell is 3cm wide and each row is 3cm tall
  Component 3 (0.20): Table has visible borders (TableGrid style)
  Component 4 (0.20): All cells have centered horizontal and vertical alignment
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_015'
FILE_PATH = f'{WORKDIR}/Desktop/game_board.docx'

# Tolerance for dimension checks: 3cm = 1701 dxa (twips), allow +-50 dxa (~0.09cm)
CM3_DXA = 1701
DXA_TOLERANCE = 50

# 3cm in EMU: 1cm = 360000 EMU
CM3_EMU = 1080000
EMU_TOLERANCE = 100000


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: 3x3 table exists with all empty cells (0.30 points) ---
    try:
        tables = doc.tables
        if len(tables) == 0:
            print("FAIL: Component 1 — no table found in document")
        else:
            table = tables[0]
            rows = table.rows
            cols = table.columns
            num_rows = len(rows)
            num_cols = len(cols)
            if num_rows == 3 and num_cols == 3:
                # Check all cells are empty
                all_empty = all(
                    cell.text.strip() == ''
                    for ri in range(3)
                    for ci in range(3)
                    for cell in [table.cell(ri, ci)]
                )
                if all_empty:
                    print("PASS: Component 1 — 3x3 table found with all empty cells (0.30 pts)")
                    total_score += 0.30
                else:
                    print("FAIL: Component 1 — 3x3 table found but some cells are not empty")
            else:
                print(f"FAIL: Component 1 — table has {num_rows} rows and {num_cols} cols, expected 3x3")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: Cell width ~3cm and row height ~3cm (0.30 points) ---
    try:
        if len(doc.tables) > 0:
            table = doc.tables[0]
            width_ok = True
            height_ok = True

            # Check cell widths in dxa
            for ri in range(min(3, len(table.rows))):
                for ci in range(min(3, len(table.columns))):
                    cell = table.cell(ri, ci)
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    cell_width_ok = False
                    if tcPr is not None:
                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is not None:
                            w_str = tcW.get(qn('w:w'))
                            w_type = tcW.get(qn('w:type'))
                            if w_str and w_type == 'dxa':
                                w_val = int(w_str)
                                if abs(w_val - CM3_DXA) <= DXA_TOLERANCE:
                                    cell_width_ok = True
                                else:
                                    print(f"FAIL: Component 2 — Cell({ri},{ci}) width={w_val} dxa, expected ~{CM3_DXA} dxa (3cm)")
                    if not cell_width_ok:
                        width_ok = False
                        break
                if not width_ok:
                    break

            # Check row heights
            for ri, row in enumerate(table.rows[:3]):
                tr = row._tr
                trPr = tr.find(qn('w:trPr'))
                row_height_ok = False
                if trPr is not None:
                    trHeight = trPr.find(qn('w:trHeight'))
                    if trHeight is not None:
                        val_str = trHeight.get(qn('w:val'))
                        if val_str:
                            val = int(val_str)
                            if abs(val - CM3_DXA) <= DXA_TOLERANCE:
                                row_height_ok = True
                            else:
                                print(f"FAIL: Component 2 — Row {ri} height={val} dxa, expected ~{CM3_DXA} dxa (3cm)")
                # Also check via EMU (row.height returns EMU)
                if not row_height_ok:
                    emu_height = row.height
                    if emu_height is not None and abs(emu_height - CM3_EMU) <= EMU_TOLERANCE:
                        row_height_ok = True
                    else:
                        print(f"FAIL: Component 2 — Row {ri} height (EMU)={emu_height}, expected ~{CM3_EMU} (3cm)")
                        height_ok = False
                        break

            if width_ok and height_ok:
                print("PASS: Component 2 — All cells are ~3cm wide and all rows are ~3cm tall (0.30 pts)")
                total_score += 0.30
            elif not width_ok:
                print("FAIL: Component 2 — One or more cell widths are not ~3cm")
            else:
                print("FAIL: Component 2 — One or more row heights are not ~3cm")
        else:
            print("FAIL: Component 2 — no table found, skipping dimension check")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: Table has visible borders (TableGrid style) (0.20 points) ---
    try:
        if len(doc.tables) > 0:
            table = doc.tables[0]
            style_name = table.style.name if table.style else None
            # TableGrid style provides visible borders on all cells
            # Also accept if explicit border XML is present on cells
            has_table_grid = style_name == 'Table Grid'

            # Fallback: check explicit border XML in tblPr or cell tcPr
            has_explicit_borders = False
            if not has_table_grid:
                tbl = table._tbl
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is not None:
                    tblBorders = tblPr.find(qn('w:tblBorders'))
                    if tblBorders is not None:
                        # Check that at least top/bottom/left/right borders exist
                        borders_found = 0
                        for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                            if tblBorders.find(qn(f'w:{side}')) is not None:
                                borders_found += 1
                        if borders_found >= 4:
                            has_explicit_borders = True

            if has_table_grid:
                print(f"PASS: Component 3 — Table uses 'Table Grid' style with visible borders (0.20 pts)")
                total_score += 0.20
            elif has_explicit_borders:
                print(f"PASS: Component 3 — Table has explicit border definitions in XML (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Table style is '{style_name}', expected 'Table Grid' or explicit borders")
        else:
            print("FAIL: Component 3 — no table found, skipping border check")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # --- Component 4: All cells have centered horizontal and vertical alignment (0.20 points) ---
    try:
        if len(doc.tables) > 0:
            table = doc.tables[0]
            all_centered = True
            failures = []

            for ri in range(min(3, len(table.rows))):
                for ci in range(min(3, len(table.columns))):
                    cell = table.cell(ri, ci)
                    tc = cell._tc

                    # Check vertical alignment
                    tcPr = tc.find(qn('w:tcPr'))
                    v_align = None
                    if tcPr is not None:
                        vAlignElem = tcPr.find(qn('w:vAlign'))
                        if vAlignElem is not None:
                            v_align = vAlignElem.get(qn('w:val'))

                    # Check horizontal alignment
                    h_align = None
                    for para in cell.paragraphs:
                        pPr = para._p.find(qn('w:pPr'))
                        if pPr is not None:
                            jc = pPr.find(qn('w:jc'))
                            if jc is not None:
                                h_align = jc.get(qn('w:val'))

                    if v_align != 'center':
                        failures.append(f"Cell({ri},{ci}) vAlign='{v_align}' (expected 'center')")
                        all_centered = False
                    if h_align != 'center':
                        failures.append(f"Cell({ri},{ci}) hAlign='{h_align}' (expected 'center')")
                        all_centered = False

            if all_centered:
                print("PASS: Component 4 — All cells have centered horizontal and vertical alignment (0.20 pts)")
                total_score += 0.20
            else:
                for f in failures[:4]:  # show at most 4 failures
                    print(f"FAIL: Component 4 — {f}")
        else:
            print("FAIL: Component 4 — no table found, skipping alignment check")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
