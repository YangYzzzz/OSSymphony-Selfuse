"""
Reward Script: Convert comma-separated names into a table
Task ID: osworld_writer_easy_016
Domain: libreoffice_writer
Scoring:
  - Component 1: Document has exactly 1 table (0.2 pts)
  - Component 2: Table has exactly 1 row and 6 columns (0.3 pts)
  - Component 3: All 6 names are correctly placed in table cells (0.4 pts)
  - Component 4: Original comma-separated text paragraph is removed (0.1 pts)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_easy_016'

# The 6 expected names in order
EXPECTED_NAMES = [
    'Alice Johnson',
    'Bob Smith',
    'Carol Davis',
    'David Lee',
    'Emma Wilson',
    'Frank Brown',
]

# The original comma-separated text that should be removed
ORIGINAL_TEXT = 'Alice Johnson, Bob Smith, Carol Davis, David Lee, Emma Wilson, Frank Brown'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Document has exactly 1 table (0.2 points)
    # This FAILS on initial (0 tables) → PASSES on golden (1 table)
    try:
        num_tables = len(doc.tables)
        if num_tables == 1:
            print(f"PASS: Component 1 — Document has exactly 1 table (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 1 — Expected 1 table, found {num_tables} tables")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table has exactly 1 row and 6 columns (0.3 points)
    # This FAILS on initial (no tables) → PASSES on golden (1 row, 6 cols)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == 1 and num_cols == 6:
                print(f"PASS: Component 2 — Table has 1 row and 6 columns (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Expected 1 row x 6 cols, found {num_rows} rows x {num_cols} cols")
        else:
            print(f"FAIL: Component 2 — No tables found, cannot check dimensions")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All 6 names are correctly placed in the table cells (0.4 points)
    # This FAILS on initial (no table cells) → PASSES on golden (correct names in cells)
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            if len(table.rows) >= 1:
                row = table.rows[0]
                cells = row.cells
                if len(cells) == 6:
                    mismatches = [
                        f"Cell {i}: expected '{expected}', found '{cell.text.strip()}'"
                        for i, (cell, expected) in enumerate(zip(cells, EXPECTED_NAMES))
                        if cell.text.strip() != expected
                    ]
                    if len(mismatches) == 0:
                        print(f"PASS: Component 3 — All 6 names correctly placed in table cells (0.4 pts)")
                        total_score += 0.4
                    else:
                        print(f"FAIL: Component 3 — Name mismatches: {mismatches}")
                else:
                    print(f"FAIL: Component 3 — Expected 6 cells in row, found {len(cells)}")
            else:
                print(f"FAIL: Component 3 — Table has no rows")
        else:
            print(f"FAIL: Component 3 — No tables found, cannot check cell names")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Original comma-separated text paragraph is removed (0.1 points)
    # This FAILS on initial (text exists) → PASSES on golden (text removed)
    try:
        all_para_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        original_still_present = any(ORIGINAL_TEXT in text for text in all_para_text)
        if not original_still_present:
            print(f"PASS: Component 4 — Original comma-separated text paragraph is removed (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 4 — Original comma-separated text still present in document")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM
file_path = f'{WORKDIR}/team_roster.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
