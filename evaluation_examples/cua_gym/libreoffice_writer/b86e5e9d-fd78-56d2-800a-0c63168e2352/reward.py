"""
Reward Script: Event Sponsorship Proposal in LibreOffice Writer
Task ID: writer_wf_087
Domain: libreoffice_writer
Scoring:
  1. Title paragraph (0.10)
  2. Audience Profile table with 5 segments (0.15)
  3. Sponsorship Tiers table with 4 tiers (0.15)
  4. Four Heading 2 tier detail sections with bullets (0.20)
  5. Media Coverage section (0.10)
  6. Previous Sponsors section (0.05)
  7. Contact Information section (0.10)
  8. Professional header with logo placeholder (0.15)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_087'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts and styles for analysis
    paras = [(p.style.name if p.style else 'Normal', p.text.strip()) for p in doc.paragraphs]
    para_texts_lower = [t.lower() for _, t in paras]

    # Component 1: Title paragraph with correct text (0.10 points)
    try:
        has_title = False
        for style, text in paras:
            if style == 'Title' and 'sponsorship proposal' in text.lower() and 'city marathon 2026' in text.lower():
                has_title = True
                break
        if has_title:
            print(f"PASS: Component 1 — Title found with 'Sponsorship Proposal' and 'City Marathon 2026' (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Title paragraph not found with required text")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Audience Profile table with 5 demographic segments (0.15 points)
    try:
        audience_table_found = False
        for table in doc.tables:
            # Check if header row contains 'Demographic' and 'Percentage'
            if len(table.rows) >= 2 and len(table.columns) >= 2:
                header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'demographic' in header_cells and 'percentage' in header_cells:
                    # Count data rows (excluding header)
                    data_rows = len(table.rows) - 1
                    if data_rows >= 5:
                        # Verify percentage values exist
                        has_percentages = False
                        for ri in range(1, len(table.rows)):
                            for cell in table.rows[ri].cells:
                                if '%' in cell.text:
                                    has_percentages = True
                                    break
                            if has_percentages:
                                break
                        if has_percentages:
                            audience_table_found = True
                            print(f"PASS: Component 2 — Audience table: {data_rows} segments with percentages (0.15 pts)")
                            total_score += 0.15
                            break
        if not audience_table_found:
            print(f"FAIL: Component 2 — Audience Profile table not found (need Demographic/Percentage headers, 5+ rows)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Sponsorship Tiers table with 4 tiers (0.15 points)
    try:
        tiers_table_found = False
        required_tiers = {'platinum', 'gold', 'silver', 'bronze'}
        for table in doc.tables:
            if len(table.rows) >= 5 and len(table.columns) >= 3:
                header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if 'tier' in header_cells and ('investment' in header_cells or 'cost' in header_cells or 'price' in header_cells):
                    # Check for the 4 required tiers
                    found_tiers = set()
                    for ri in range(1, len(table.rows)):
                        first_cell = table.rows[ri].cells[0].text.strip().lower()
                        for tier in required_tiers:
                            if tier in first_cell:
                                found_tiers.add(tier)
                    if found_tiers == required_tiers:
                        tiers_table_found = True
                        print(f"PASS: Component 3 — Sponsorship Tiers table with all 4 tiers found (0.15 pts)")
                        total_score += 0.15
                        break
        if not tiers_table_found:
            print(f"FAIL: Component 3 — Sponsorship Tiers table not found (need Tier/Investment headers, 4 tiers)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Four Heading 2 tier detail sections with bullet lists (0.20 points)
    try:
        # Find Heading 2 paragraphs that reference tier names
        tier_keywords = ['platinum', 'gold', 'silver', 'bronze']
        heading2_tiers = set()
        tier_has_bullets = {}

        for i, (style, text) in enumerate(paras):
            if style == 'Heading 2':
                text_lower = text.lower()
                for kw in tier_keywords:
                    if kw in text_lower:
                        heading2_tiers.add(kw)
                        # Check if followed by bullet items
                        bullet_count = 0
                        for j in range(i + 1, len(paras)):
                            if paras[j][0] in ('List Bullet', 'List Bullet 2', 'List Bullet 3'):
                                bullet_count += 1
                            elif paras[j][0].startswith('Heading'):
                                break
                            elif paras[j][1] == '':
                                continue
                            else:
                                break
                        tier_has_bullets[kw] = bullet_count > 0

        # Score: 0.05 per tier section with bullets
        tier_detail_score = 0.0
        for kw in tier_keywords:
            if kw in heading2_tiers and tier_has_bullets.get(kw, False):
                tier_detail_score += 0.05

        if tier_detail_score > 0:
            found_count = sum(1 for kw in tier_keywords if kw in heading2_tiers and tier_has_bullets.get(kw, False))
            print(f"PASS: Component 4 — {found_count}/4 tier detail sections with bullets ({tier_detail_score:.2f} pts)")
            total_score += tier_detail_score
        else:
            print(f"FAIL: Component 4 — No Heading 2 tier detail sections with bullet lists found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Media Coverage section (0.10 points)
    try:
        media_section_found = False
        for i, (style, text) in enumerate(paras):
            if style.startswith('Heading') and 'media' in text.lower() and 'coverage' in text.lower():
                # Check for content after this heading
                has_content = False
                for j in range(i + 1, min(i + 15, len(paras))):
                    if paras[j][0].startswith('Heading'):
                        break
                    if paras[j][1]:
                        has_content = True
                        break
                if has_content:
                    media_section_found = True
                    break
        if media_section_found:
            print(f"PASS: Component 5 — Media Coverage section found with content (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — Media Coverage section not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Previous Sponsors section (0.05 points)
    try:
        sponsors_section_found = False
        for i, (style, text) in enumerate(paras):
            if style.startswith('Heading') and 'sponsor' in text.lower() and ('previous' in text.lower() or 'past' in text.lower()):
                # Check for content (list items or paragraphs)
                has_content = False
                for j in range(i + 1, min(i + 15, len(paras))):
                    if paras[j][0].startswith('Heading'):
                        break
                    if paras[j][1]:
                        has_content = True
                        break
                if has_content:
                    sponsors_section_found = True
                    break
        if sponsors_section_found:
            print(f"PASS: Component 6 — Previous Sponsors section found with content (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — Previous Sponsors section not found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Contact Information section (0.10 points)
    try:
        contact_section_found = False
        for i, (style, text) in enumerate(paras):
            if style.startswith('Heading') and 'contact' in text.lower():
                # Check for contact details after this heading
                has_contact_info = False
                for j in range(i + 1, min(i + 15, len(paras))):
                    if paras[j][0].startswith('Heading'):
                        break
                    t = paras[j][1].lower()
                    if 'email' in t or 'phone' in t or '@' in t or 'name:' in t:
                        has_contact_info = True
                        break
                if has_contact_info:
                    contact_section_found = True
                    break
        if contact_section_found:
            print(f"PASS: Component 7 — Contact Information section found with details (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 — Contact Information section not found with contact details")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    # Component 8: Professional header with event logo placeholder (0.15 points)
    try:
        header_found = False
        for section in doc.sections:
            header = section.header
            if header and header.paragraphs:
                header_text = ' '.join(p.text for p in header.paragraphs).strip()
                if header_text:
                    # Check for logo placeholder reference and event name
                    has_logo_ref = 'logo' in header_text.lower() or '[' in header_text
                    has_event_ref = 'marathon' in header_text.lower() or 'sponsorship' in header_text.lower()
                    if has_logo_ref and has_event_ref:
                        header_found = True
                        print(f"PASS: Component 8 — Header found: '{header_text[:80]}' (0.15 pts)")
                        total_score += 0.15
                        break
                    elif header_text:
                        # Partial: header exists but may not have logo placeholder
                        header_found = True
                        print(f"PARTIAL: Component 8 — Header exists but missing logo placeholder or event ref (0.08 pts)")
                        total_score += 0.08
                        break
        if not header_found:
            print(f"FAIL: Component 8 — No header found or header is empty")
    except Exception as e:
        print(f"ERROR: Component 8 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
