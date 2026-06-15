"""
Reward Script: Delete the word 'very' every time it appears in the document.
Task ID: writer_edit_058
Domain: libreoffice_writer
Scoring:
  Component 1 (0.7): All occurrences of 'very' (word boundary match) are removed
  Component 2 (0.3): Document text integrity is preserved (correct paragraph count,
                      key phrases intact without 'very', and shorter total length)
"""

import os
import re

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_edit_058'
FILE_NAME = 'student_essay.docx'


def get_full_text(doc):
    """Return full document text by joining all paragraph texts."""
    return ' '.join(para.text for para in doc.paragraphs)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — precondition gate
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get full document text
    try:
        full_text = get_full_text(doc)
    except Exception as e:
        print(f"CRITICAL: Cannot extract text from document: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 12 occurrences of 'very' have been removed (0.7 points)
    # The task requires deleting 'very ' (very followed by space) 12 times.
    # We verify using word-boundary regex to catch all occurrences.
    # This FAILS on initial_env (12 occurrences) and PASSES on golden_env (0 occurrences).
    try:
        very_occurrences = re.findall(r'\bvery\b', full_text, re.IGNORECASE)
        count = len(very_occurrences)
        if count == 0:
            print(f"PASS: Component 1 — No occurrences of 'very' found in document (0.7 pts)")
            total_score += 0.7
        else:
            print(f"FAIL: Component 1 — Found {count} occurrence(s) of 'very' remaining: {very_occurrences[:5]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Document integrity preserved — paragraph count and key phrases intact (0.3 points)
    # The document should still have 27 paragraphs and contain characteristic phrases
    # (e.g. 'important topic', 'serious challenge', etc.) indicating text was not corrupted.
    # The text length should be shorter than the original (~4854) but at least 4700 chars
    # (accounting for removal of up to 12 x "very " = 60 chars).
    # This FAILS on initial_env (because the 'very' check in Component 1 must fail first
    # but this checks structural integrity separately — note this is designed to be a
    # sub-condition confirming non-corruption after the edit).
    # The key sub-check here is that the phrases EXIST (which requires Component 1 to pass
    # for phrases like 'very important' → 'important', so without 'very' the phrase 'important topic'
    # should still be present).
    try:
        num_paras = len(doc.paragraphs)
        # Key phrases that should remain in the document after removing 'very'
        # These are phrases that existed in context and remain after deletion
        key_phrases = [
            'important topic',
            'serious challenge',
            'alarming statistic',
            'rapid pace',
            'positive results'
        ]
        phrases_found = [p for p in key_phrases if p.lower() in full_text.lower()]
        text_length = len(full_text)

        # Integrity conditions:
        # 1. Paragraph count should be 27 (unchanged)
        # 2. At least 4 of 5 key phrases found
        # 3. Total text length should be in range [4700, 4854)
        #    (shorter than initial 4854 due to removal, but not empty/corrupted)
        para_ok = (num_paras == 27)
        phrases_ok = (len(phrases_found) >= 4)
        length_ok = (4700 <= text_length < 4854)

        integrity_ok = (para_ok and phrases_ok and length_ok)
        if integrity_ok:
            print(f"PASS: Component 2 — Document integrity preserved: {num_paras} paragraphs, {len(phrases_found)}/{len(key_phrases)} key phrases found, text length={text_length} (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Document integrity issue: paragraphs={num_paras} (expected 27, ok={para_ok}), phrases found={len(phrases_found)}/{len(key_phrases)} (ok={phrases_ok}), text length={text_length} (ok={length_ok})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path
file_path = f'{WORKDIR}/{FILE_NAME}'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
