"""
Initial Setup: Student essay that overuses the word 'very' (12 occurrences of 'very ')
Task ID: writer_edit_058
Domain: libreoffice_writer
Note: Essay intentionally avoids the word 'every' to prevent false matches with 'very '.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'student_essay'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure the Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('The Importance of Environmental Conservation', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author and date line
    author_para = doc.add_paragraph('Emily Rodriguez | Environmental Studies 101 | March 2025')
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # blank line

    # --- INTRODUCTION (Page 1) ---
    # 'very important' [1]
    intro_heading = doc.add_heading('Introduction', level=2)

    para1 = doc.add_paragraph(
        'Environmental conservation has become a very important topic in recent decades. '
        'As industrialization expands, the natural world faces threats that were once unimaginable. '
        'Scientists, policymakers, and ordinary citizens must work together to protect ecosystems '
        'that have taken millions of years to develop.'
    )
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 'very large' [2], 'very significant' [3]
    para2 = doc.add_paragraph(
        'The scale of environmental degradation is very large and touches all corners of the globe. '
        'From the melting ice caps in the Arctic to the disappearing rainforests of the Amazon basin, '
        'the evidence of climate change and habitat destruction is very significant and impossible to ignore. '
        'This essay examines key dimensions of the environmental crisis and proposes pathways toward '
        'meaningful solutions.'
    )
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 'very essential' [4]
    para3 = doc.add_paragraph(
        'Understanding the root causes of environmental damage is very essential before solutions can '
        'be designed effectively. Human activities—ranging from industrial agriculture to urban expansion—have '
        'disrupted natural cycles that once maintained balance in global ecosystems.'
    )
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- BODY SECTION 1 (Page 2) ---
    doc.add_page_break()

    section1_heading = doc.add_heading('The Impact of Pollution', level=2)

    # 'very serious' [5]
    para4 = doc.add_paragraph(
        'Pollution, in its many forms, represents a very serious challenge for both human health and '
        'natural habitats. Air pollution from factories and vehicles releases greenhouse gases that trap '
        'heat in the atmosphere, contributing directly to climate change. Water pollution from agricultural '
        'runoff and industrial discharge has rendered many rivers and lakes unable to support aquatic life.'
    )
    para4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 'very severe' [6], 'very alarming' [7]
    para5 = doc.add_paragraph(
        'The consequences for biodiversity have been very severe. Coral reefs, which are among the most '
        'biologically diverse ecosystems on Earth, have experienced widespread bleaching caused by '
        'rising ocean temperatures. Marine biologists report that over 50 percent of coral coverage has '
        'been lost globally since the 1970s, which is a very alarming statistic that demands urgent action.'
    )
    para5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 'very challenging' [8]
    para6 = doc.add_paragraph(
        'Soil contamination presents another very challenging dimension of the pollution crisis. '
        'Heavy metals and persistent organic pollutants accumulate in agricultural soils, reducing fertility '
        'and entering the food chain through crops. Communities near industrial sites often suffer '
        'disproportionately from exposure to toxic substances.'
    )
    para6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- BODY SECTION 2 ---
    section2_heading = doc.add_heading('Deforestation and Habitat Loss', level=2)

    # 'very rapid' [9], 'very crucial' [10]
    para7 = doc.add_paragraph(
        'Deforestation continues at a very rapid pace, particularly in tropical regions where biodiversity '
        'is highest. Each year, millions of hectares of forest are cleared for agriculture, cattle ranching, '
        'and timber extraction. These forests are very crucial carbon sinks that absorb atmospheric CO2; '
        'their removal accelerates the greenhouse effect dramatically.'
    )
    para7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 'very deeply' [11]
    para8 = doc.add_paragraph(
        'Indigenous communities that have coexisted with these forests for centuries are also very deeply '
        'affected by deforestation. Their traditional knowledge, cultural practices, and livelihoods depend '
        'on intact forest ecosystems. When forests disappear, entire ways of life are lost irreversibly. '
        'International organizations have called for legally binding agreements to protect both forests '
        'and indigenous land rights.'
    )
    para8.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- BODY SECTION 3 (Page 3) ---
    doc.add_page_break()

    section3_heading = doc.add_heading('Pathways to Conservation', level=2)

    para9 = doc.add_paragraph(
        'Addressing the environmental crisis requires solutions that are coordinated across national borders. '
        'Renewable energy technologies, including solar and wind power, have become increasingly affordable '
        'and represent a promising alternative to fossil fuels. Investment in green infrastructure can '
        'stimulate economic growth while reducing carbon emissions.'
    )
    para9.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para10 = doc.add_paragraph(
        'Education and public awareness campaigns play a vital role in building the social support '
        'necessary for policy change. When citizens understand the connection between their daily choices '
        'and environmental outcomes, behavioral shifts can occur at scale. Schools, media, and community '
        'organizations all have responsibilities in spreading environmental literacy.'
    )
    para10.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # --- CONCLUSION ---
    conclusion_heading = doc.add_heading('Conclusion', level=2)

    # 'very positive' [12]
    para11 = doc.add_paragraph(
        'The environmental challenges we face are formidable, but not insurmountable. Decades of '
        'conservation research have demonstrated that ecosystems can recover when given the opportunity. '
        'Protected areas, pollution regulations, and sustainable land-use practices have all shown '
        'very positive results in regions where they have been consistently applied.'
    )
    para11.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para12 = doc.add_paragraph(
        'Ultimately, the future of environmental conservation depends on collective willpower and '
        'long-term commitment. The decisions made by governments, corporations, and individuals over '
        'the next few decades will determine whether the planet\'s biodiversity and natural systems '
        'can be preserved for future generations. The time to act is now, and the stakes are high indeed.'
    )
    para12.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # References section
    doc.add_paragraph('')
    ref_heading = doc.add_heading('References', level=2)

    doc.add_paragraph(
        'Hansen, J., & Sato, M. (2024). Climate sensitivity, sea level, and atmospheric CO2. '
        'Journal of Environmental Science, 58(3), 112-134.'
    )
    doc.add_paragraph(
        'Rodriguez, C., & Patel, A. (2023). Deforestation rates in tropical biomes: A decade review. '
        'Conservation Biology, 37(2), 445-460.'
    )
    doc.add_paragraph(
        'World Wildlife Fund. (2025). Living Planet Report 2025. WWF International, Gland, Switzerland.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count 'very ' occurrences to confirm (should be exactly 12)
    doc_check = Document(OUTPUT)
    count = 0
    for para in doc_check.paragraphs:
        count += para.text.count('very ')
    print(f'Occurrences of "very " in document: {count} (expected: 12)')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
