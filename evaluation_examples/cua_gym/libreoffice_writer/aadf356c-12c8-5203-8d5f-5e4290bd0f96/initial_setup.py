"""
Initial Setup: Create compliance document with 22 tracked changes
Task ID: writer_rm_023
Domain: libreoffice_writer

Creates a compliance document with tracked changes from two authors:
- External_Reviewer: 12 changes (terminology updates, legal phrasing)
- Intern_Jones: 10 changes (informal language, unnecessary additions)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_023'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'

def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def make_run_element(text, bold=False, italic=False, font_name="Calibri", font_size_pt=11):
    """Create a w:r element with text and optional formatting."""
    r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{text}</w:t></w:r>')
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    if bold:
        rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
    if italic:
        rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
    if font_name:
        rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>'))
    if font_size_pt:
        half_pt = font_size_pt * 2
        rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{half_pt}"/>'))
    r.insert(0, rPr)
    return r


def make_ins(run_elem, author, rev_id, date="2026-03-28T10:00:00Z"):
    """Wrap a run element in w:ins (tracked insertion)."""
    ins = parse_xml(
        f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author}" w:date="{date}"/>'
    )
    ins.append(run_elem)
    return ins


def make_del(text, author, rev_id, date="2026-03-28T10:00:00Z",
             bold=False, italic=False, font_name="Calibri", font_size_pt=11):
    """Create a w:del element (tracked deletion) containing w:r with w:delText."""
    r = parse_xml(f'<w:r {nsdecls("w")}><w:delText xml:space="preserve">{text}</w:delText></w:r>')
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    if bold:
        rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
    if italic:
        rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
    if font_name:
        rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>'))
    if font_size_pt:
        half_pt = font_size_pt * 2
        rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{half_pt}"/>'))
    r.insert(0, rPr)
    d = parse_xml(
        f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author}" w:date="{date}"/>'
    )
    d.append(r)
    return d


def add_plain_run(para_elem, text, bold=False, italic=False, font_name="Calibri", font_size_pt=11):
    """Add a plain run to a paragraph element."""
    r = make_run_element(text, bold=bold, italic=italic, font_name=font_name, font_size_pt=font_size_pt)
    para_elem.append(r)
    return r


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Corporate Compliance Policy Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.add_run('Regulatory Affairs Division — Version 4.2')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()  # spacer

    # Section 1 heading
    doc.add_heading('1. Purpose and Scope', level=1)

    # --- Build paragraphs with tracked changes via XML manipulation ---
    # We'll use the document body element to append paragraphs with revision marks

    body = doc.element.body
    rev_id = 100  # starting revision ID

    # ============================================================
    # PARAGRAPH: Section 1 content with External_Reviewer change #1
    # Change: "rules" -> "regulatory requirements" (delete "rules", insert "regulatory requirements")
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'This document establishes the framework for all corporate compliance ')
    p.append(make_del('rules', 'External_Reviewer', rev_id, "2026-03-25T09:15:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('regulatory requirements'), 'External_Reviewer', rev_id, "2026-03-25T09:15:00Z"))
    rev_id += 1
    add_plain_run(p, ' applicable to Meridian Global Holdings and all subsidiary entities operating within domestic and international jurisdictions.')
    body.append(p)

    # ============================================================
    # PARAGRAPH: With Intern_Jones change #1
    # Intern adds informal text: inserts "basically" before "ensures"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'The scope of this policy ')
    p.append(make_ins(make_run_element('basically '), 'Intern_Jones', rev_id, "2026-03-26T14:20:00Z"))
    rev_id += 1
    add_plain_run(p, 'ensures adherence to all federal, state, and local statutes governing financial reporting, data protection, and workplace safety standards.')
    body.append(p)

    # Section 2
    h2 = doc.add_heading('2. Definitions and Key Terms', level=1)

    # ============================================================
    # External_Reviewer change #2: "must follow" -> "shall comply with"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'All personnel, contractors, and third-party vendors ')
    p.append(make_del('must follow', 'External_Reviewer', rev_id, "2026-03-25T09:30:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('shall comply with'), 'External_Reviewer', rev_id, "2026-03-25T09:30:00Z"))
    rev_id += 1
    add_plain_run(p, ' the provisions outlined herein. Non-compliance may result in disciplinary action up to and including termination of employment or contractual obligations.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #2: replaces "disciplinary action" with "getting in trouble"
    # (delete "disciplinary action" insert "getting in trouble")
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Violations identified during internal audits will be subject to ')
    p.append(make_del('disciplinary review', 'Intern_Jones', rev_id, "2026-03-26T14:35:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('getting in trouble'), 'Intern_Jones', rev_id, "2026-03-26T14:35:00Z"))
    rev_id += 1
    add_plain_run(p, ' and escalation procedures as defined in Section 7.')
    body.append(p)

    # Section 3
    doc.add_heading('3. Regulatory Framework', level=1)

    # ============================================================
    # External_Reviewer change #3: "laws" -> "applicable statutes and regulations"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Meridian Global Holdings operates under the oversight of multiple regulatory bodies. All business operations must conform to ')
    p.append(make_del('laws', 'External_Reviewer', rev_id, "2026-03-25T10:00:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('applicable statutes and regulations'), 'External_Reviewer', rev_id, "2026-03-25T10:00:00Z"))
    rev_id += 1
    add_plain_run(p, ' including but not limited to the Sarbanes-Oxley Act, the Dodd-Frank Wall Street Reform Act, and the General Data Protection Regulation.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #3: inserts unnecessary sentence
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Each department head is responsible for ensuring their team understands the regulatory obligations specific to their function.')
    p.append(make_ins(make_run_element(' This is pretty important stuff that everyone should pay attention to.'), 'Intern_Jones', rev_id, "2026-03-26T15:00:00Z"))
    rev_id += 1
    body.append(p)

    # ============================================================
    # External_Reviewer change #4: "check" -> "conduct periodic assessments of"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'The Compliance Officer shall ')
    p.append(make_del('check', 'External_Reviewer', rev_id, "2026-03-25T10:15:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('conduct periodic assessments of'), 'External_Reviewer', rev_id, "2026-03-25T10:15:00Z"))
    rev_id += 1
    add_plain_run(p, ' all departmental compliance records on a quarterly basis and report findings to the Board of Directors.')
    body.append(p)

    # Section 4
    doc.add_heading('4. Data Protection and Privacy', level=1)

    # ============================================================
    # External_Reviewer change #5: "keep safe" -> "safeguard the confidentiality of"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'All employees are required to ')
    p.append(make_del('keep safe', 'External_Reviewer', rev_id, "2026-03-25T10:30:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('safeguard the confidentiality of'), 'External_Reviewer', rev_id, "2026-03-25T10:30:00Z"))
    rev_id += 1
    add_plain_run(p, ' personally identifiable information (PII) and protected health information (PHI) in accordance with HIPAA and GDPR provisions.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #4: "personally identifiable information" -> "personal info"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Data breach notification protocols require immediate reporting to the ')
    p.append(make_del('Chief Information Security Officer', 'Intern_Jones', rev_id, "2026-03-26T15:15:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('IT guy'), 'Intern_Jones', rev_id, "2026-03-26T15:15:00Z"))
    rev_id += 1
    add_plain_run(p, ' within 24 hours of discovery, followed by a comprehensive incident report within 72 hours.')
    body.append(p)

    # ============================================================
    # External_Reviewer change #6: "broken" -> "compromised"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Any system determined to be ')
    p.append(make_del('broken', 'External_Reviewer', rev_id, "2026-03-25T11:00:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('compromised'), 'External_Reviewer', rev_id, "2026-03-25T11:00:00Z"))
    rev_id += 1
    add_plain_run(p, ' must be isolated from the network immediately and subjected to forensic analysis before reconnection is authorized.')
    body.append(p)

    # Section 5
    doc.add_heading('5. Financial Reporting Standards', level=1)

    # ============================================================
    # Intern_Jones change #5: inserts "honestly" (unnecessary addition)
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'All financial statements must ')
    p.append(make_ins(make_run_element('honestly '), 'Intern_Jones', rev_id, "2026-03-26T15:30:00Z"))
    rev_id += 1
    add_plain_run(p, 'accurately reflect the fiscal position of the organization in accordance with Generally Accepted Accounting Principles (GAAP) and International Financial Reporting Standards (IFRS).')
    body.append(p)

    # ============================================================
    # External_Reviewer change #7: "money handling" -> "financial transaction processing"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Internal controls for ')
    p.append(make_del('money handling', 'External_Reviewer', rev_id, "2026-03-25T11:15:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('financial transaction processing'), 'External_Reviewer', rev_id, "2026-03-25T11:15:00Z"))
    rev_id += 1
    add_plain_run(p, ' must include segregation of duties, dual authorization for transactions exceeding $50,000, and automated reconciliation of all ledger entries.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #6: "segregation of duties" -> "splitting up the work"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Quarterly audits shall examine the effectiveness of ')
    p.append(make_del('risk mitigation strategies', 'Intern_Jones', rev_id, "2026-03-26T15:45:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('ways we deal with problems'), 'Intern_Jones', rev_id, "2026-03-26T15:45:00Z"))
    rev_id += 1
    add_plain_run(p, ' and ensure alignment with the enterprise risk management framework established by the Risk Committee.')
    body.append(p)

    # Section 6
    doc.add_heading('6. Anti-Corruption and Bribery Prevention', level=1)

    # ============================================================
    # External_Reviewer change #8: "gifts" -> "gratuities, hospitality, or other inducements"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'No employee shall offer, solicit, or accept ')
    p.append(make_del('gifts', 'External_Reviewer', rev_id, "2026-03-25T11:30:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('gratuities, hospitality, or other inducements'), 'External_Reviewer', rev_id, "2026-03-25T11:30:00Z"))
    rev_id += 1
    add_plain_run(p, ' that could reasonably be perceived as an attempt to influence business decisions or secure preferential treatment.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #7: inserts "like free lunches or whatever"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'All business entertainment expenses')
    p.append(make_ins(make_run_element(', like free lunches or whatever,'), 'Intern_Jones', rev_id, "2026-03-26T16:00:00Z"))
    rev_id += 1
    add_plain_run(p, ' must be pre-approved by the department manager and documented with itemized receipts within five business days of the expenditure.')
    body.append(p)

    # ============================================================
    # External_Reviewer change #9: "Foreign Corrupt Practices Act" terminology fix
    # "the anti-bribery law" -> "the Foreign Corrupt Practices Act (FCPA)"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'International operations must comply with ')
    p.append(make_del('the anti-bribery law', 'External_Reviewer', rev_id, "2026-03-25T11:45:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('the Foreign Corrupt Practices Act (FCPA)'), 'External_Reviewer', rev_id, "2026-03-25T11:45:00Z"))
    rev_id += 1
    add_plain_run(p, ' and the UK Bribery Act 2010, regardless of the jurisdiction in which the transaction occurs.')
    body.append(p)

    # Section 7
    doc.add_heading('7. Whistleblower Protection', level=1)

    # ============================================================
    # External_Reviewer change #10: "tell on" -> "report suspected violations through"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Employees who ')
    p.append(make_del('tell on', 'External_Reviewer', rev_id, "2026-03-25T12:00:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('report suspected violations through'), 'External_Reviewer', rev_id, "2026-03-25T12:00:00Z"))
    rev_id += 1
    add_plain_run(p, ' the confidential ethics hotline or other designated channels are protected from retaliation under federal and state whistleblower statutes.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #8: "retaliation" -> "being treated badly"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'The Ethics Committee shall investigate all reports of ')
    p.append(make_del('retaliatory conduct', 'Intern_Jones', rev_id, "2026-03-26T16:15:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('people being mean to whistleblowers'), 'Intern_Jones', rev_id, "2026-03-26T16:15:00Z"))
    rev_id += 1
    add_plain_run(p, ' and take appropriate corrective action within 30 calendar days of the initial report.')
    body.append(p)

    # Section 8
    doc.add_heading('8. Training and Certification', level=1)

    # ============================================================
    # External_Reviewer change #11: "training" -> "mandatory compliance education programs"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'All new hires must complete ')
    p.append(make_del('training', 'External_Reviewer', rev_id, "2026-03-25T12:15:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('mandatory compliance education programs'), 'External_Reviewer', rev_id, "2026-03-25T12:15:00Z"))
    rev_id += 1
    add_plain_run(p, ' within 90 days of their start date. Annual recertification is required for all employees in regulated business units.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #9: inserts "I think this is a lot of training lol"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Training modules shall cover anti-money laundering protocols, insider trading prohibitions, conflict of interest disclosures, and workplace harassment prevention.')
    p.append(make_ins(make_run_element(' I think this is a lot of training lol.'), 'Intern_Jones', rev_id, "2026-03-26T16:30:00Z"))
    rev_id += 1
    body.append(p)

    # Section 9
    doc.add_heading('9. Enforcement and Penalties', level=1)

    # ============================================================
    # External_Reviewer change #12: "punishment" -> "proportionate sanctions"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'Violations of this policy shall result in ')
    p.append(make_del('punishment', 'External_Reviewer', rev_id, "2026-03-25T12:30:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('proportionate sanctions'), 'External_Reviewer', rev_id, "2026-03-25T12:30:00Z"))
    rev_id += 1
    add_plain_run(p, ' determined by the severity and recurrence of the infraction, as adjudicated by the Compliance Review Board.')
    body.append(p)

    # ============================================================
    # Intern_Jones change #10: "adjudicated" -> "decided"
    # ============================================================
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:pPr></w:p>')
    add_plain_run(p, 'The Compliance Review Board reserves the right to refer matters of ')
    p.append(make_del('criminal misconduct', 'Intern_Jones', rev_id, "2026-03-26T16:45:00Z"))
    rev_id += 1
    p.append(make_ins(make_run_element('really bad stuff'), 'Intern_Jones', rev_id, "2026-03-26T16:45:00Z"))
    rev_id += 1
    add_plain_run(p, ' to the appropriate law enforcement authorities and regulatory agencies for further investigation and prosecution.')
    body.append(p)

    # Final paragraph (no changes)
    p_final = doc.add_paragraph()
    p_final.add_run('This policy is effective as of January 1, 2026, and supersedes all prior versions. Review and updates shall occur annually or as necessitated by changes in applicable law.').font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
