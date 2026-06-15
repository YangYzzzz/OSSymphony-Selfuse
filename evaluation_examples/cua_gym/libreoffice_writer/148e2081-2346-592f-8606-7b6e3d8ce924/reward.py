"""
Reward Script: Business Continuity Plan Summary in LibreOffice Writer
Task ID: writer_wf_053
Domain: libreoffice_writer
Scoring:
  C1 (0.10) - Title paragraph with correct text
  C2 (0.20) - 6 Heading 1 sections with required names
  C3 (0.15) - 3 Heading 2 recovery phase subsections
  C4 (0.20) - Risk Assessment table (header + 5 data rows, 4 cols)
  C5 (0.15) - Emergency Contacts table (header + 4 data rows, 4 cols)
  C6 (0.10) - CONFIDENTIAL text in header
  C7 (0.10) - Table of Contents section present
"""

import os


WORKDIR = '/home/user'
TASK_ID = 'writer_wf_053'


def persist_app_state(domain):
    """Try to save any open LibreOffice document via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect paragraph info
    paragraphs = doc.paragraphs
    headings_h1 = []
    headings_h2 = []
    title_para = None

    for p in paragraphs:
        sname = p.style.name if p.style else ''
        if sname == 'Title':
            title_para = p
        elif sname == 'Heading 1':
            headings_h1.append(p.text.strip())
        elif sname == 'Heading 2':
            headings_h2.append(p.text.strip())

    # Component 1: Title paragraph (0.10 points)
    try:
        if title_para and 'business continuity plan' in title_para.text.lower() and 'nexora' in title_para.text.lower():
            print(f"PASS: Component 1 — Title found: '{title_para.text[:80]}' (0.10 pts)")
            total_score += 0.10
        else:
            found_text = title_para.text[:80] if title_para else 'None'
            print(f"FAIL: Component 1 — Expected title with 'Business Continuity Plan' and 'Nexora', found: '{found_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 6 Heading 1 sections with required names (0.20 points)
    # Required sections: Purpose and Scope, Risk Assessment, Emergency Contacts,
    # Recovery Procedures, Communication Plan, Plan Maintenance
    try:
        required_h1 = [
            'purpose and scope',
            'risk assessment',
            'emergency contacts',
            'recovery procedures',
            'communication plan',
            'plan maintenance',
        ]
        h1_lower = [h.lower() for h in headings_h1]
        matched = sum(1 for req in required_h1 if any(req in h for h in h1_lower))
        if matched >= 6:
            print(f"PASS: Component 2 — All 6 required H1 sections found (0.20 pts)")
            total_score += 0.20
        elif matched >= 4:
            partial = round(0.20 * (matched / 6), 2)
            print(f"PARTIAL: Component 2 — {matched}/6 H1 sections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched}/6 required H1 sections. Found H1s: {headings_h1}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 3 Heading 2 recovery phase subsections (0.15 points)
    try:
        phase_keywords = ['phase 1', 'phase 2', 'phase 3']
        h2_lower = [h.lower() for h in headings_h2]
        phase_matches = sum(1 for pk in phase_keywords if any(pk in h for h in h2_lower))
        if phase_matches >= 3:
            print(f"PASS: Component 3 — All 3 recovery phase H2 subsections found (0.15 pts)")
            total_score += 0.15
        elif phase_matches >= 1:
            partial = round(0.15 * (phase_matches / 3), 2)
            print(f"PARTIAL: Component 3 — {phase_matches}/3 phase H2 subsections found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No recovery phase H2 subsections found. H2s: {headings_h2}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Risk Assessment table (0.20 points)
    # Should have header row (Risk, Probability, Impact, Mitigation) + 5 data rows = 6 total rows, 4 cols
    try:
        risk_table = None
        for table in doc.tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'risk' in header_cells and ('probability' in header_cells or 'likelihood' in header_cells):
                risk_table = table
                break

        if risk_table is not None:
            num_rows = len(risk_table.rows)
            num_cols = len(risk_table.rows[0].cells)
            data_rows = num_rows - 1  # subtract header

            # Check columns: need at least 4 (Risk, Probability, Impact, Mitigation)
            col_ok = num_cols >= 4
            # Check data rows: need at least 5
            rows_ok = data_rows >= 5
            # Check that data cells are non-empty
            non_empty_data = 0
            for r_idx in range(1, min(num_rows, 6)):
                if risk_table.rows[r_idx].cells[0].text.strip():
                    non_empty_data += 1

            if col_ok and rows_ok and non_empty_data >= 5:
                print(f"PASS: Component 4 — Risk table found: {num_cols} cols, {data_rows} data rows, {non_empty_data} non-empty (0.20 pts)")
                total_score += 0.20
            elif col_ok and data_rows >= 3 and non_empty_data >= 3:
                partial = 0.10
                print(f"PARTIAL: Component 4 — Risk table found but incomplete: {data_rows} data rows, {non_empty_data} non-empty ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Risk table structure invalid: {num_cols} cols, {data_rows} data rows")
        else:
            print(f"FAIL: Component 4 — No risk assessment table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Emergency Contacts table (0.15 points)
    # Should have header row (Role, Name, Phone, Email) + 4 data rows = 5 total rows, 4 cols
    try:
        contacts_table = None
        for table in doc.tables:
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'role' in header_cells and ('name' in header_cells or 'contact' in header_cells):
                contacts_table = table
                break

        if contacts_table is not None:
            num_rows = len(contacts_table.rows)
            num_cols = len(contacts_table.rows[0].cells)
            data_rows = num_rows - 1

            col_ok = num_cols >= 4
            rows_ok = data_rows >= 4
            non_empty_data = 0
            for r_idx in range(1, min(num_rows, 5)):
                if contacts_table.rows[r_idx].cells[0].text.strip():
                    non_empty_data += 1

            if col_ok and rows_ok and non_empty_data >= 4:
                print(f"PASS: Component 5 — Contacts table found: {num_cols} cols, {data_rows} data rows, {non_empty_data} non-empty (0.15 pts)")
                total_score += 0.15
            elif col_ok and data_rows >= 2 and non_empty_data >= 2:
                partial = 0.08
                print(f"PARTIAL: Component 5 — Contacts table found but incomplete: {data_rows} data rows ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — Contacts table structure invalid: {num_cols} cols, {data_rows} data rows")
        else:
            print(f"FAIL: Component 5 — No emergency contacts table found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: CONFIDENTIAL header text (0.10 points)
    try:
        header_has_confidential = False
        for section in doc.sections:
            header = section.header
            if header and header.paragraphs:
                for hp in header.paragraphs:
                    if 'confidential' in hp.text.lower():
                        header_has_confidential = True
                        break

        if header_has_confidential:
            print(f"PASS: Component 6 — 'CONFIDENTIAL' found in header (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 6 — 'CONFIDENTIAL' not found in document header")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Table of Contents section (0.10 points)
    # Check for a Heading 1 or mention of TOC
    try:
        has_toc = False
        for p in paragraphs:
            sname = p.style.name if p.style else ''
            txt = p.text.lower().strip()
            # Check for "table of contents" heading or TOC entries
            if 'table of contents' in txt and sname in ('Heading 1', 'Heading 2', 'TOC Heading', 'Title'):
                has_toc = True
                break

        # Also check if there are TOC field codes in the XML
        if not has_toc:
            from docx.oxml.ns import qn
            body = doc.element.body
            toc_fields = body.findall('.//' + qn('w:instrText'))
            for f in toc_fields:
                if f.text and 'TOC' in f.text.upper():
                    has_toc = True
                    break

        # Also check for a heading that says "Table of Contents" even with just Heading 1
        if not has_toc:
            if 'table of contents' in [h.lower() for h in headings_h1]:
                has_toc = True

        if has_toc:
            print(f"PASS: Component 7 — Table of Contents section found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 7 — No Table of Contents section found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
