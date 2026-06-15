"""
Reward Script: Change the table width to be exactly 80% of the page width (relative width).
Task ID: writer_tbl_024
Domain: libreoffice_writer
Scoring:
  Component 1: Table width type is 'pct' (relative/percentage) — 0.5 pts
  Component 2: Table width value is 4000 (= 80% in OOXML fiftieths-of-a-percent) — 0.3 pts
  Component 3: Cell contents are unchanged (data integrity compound check) — 0.2 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_024'

# Expected values from task ground truth:
# In OOXML, tblW with type='pct' uses fiftieths of a percent.
# 80% = 80 * 50 = 4000 fiftieths-of-a-percent
EXPECTED_TYPE = 'pct'
EXPECTED_W = 4000       # 4000 / 50 = 80.0%
TOLERANCE = 50          # Allow ±1% tolerance (±50 in fiftieths)

# Expected cell contents (unchanged from initial state)
EXPECTED_CELLS = [
    ['Parameter', 'Value'],
    ['Temperature', '22\u00b0C'],
    ['Humidity', '45%'],
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if not doc.tables:
        print("FAIL: Document contains no tables")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))

    if tblPr is None:
        print("FAIL: Table has no tblPr (table properties element)")
        print("REWARD: 0.0")
        return 0.0

    tblW = tblPr.find(qn('w:tblW'))

    # Component 1: Table width type is 'pct' (relative width) — 0.5 points
    # This FAILS on initial_env (type='auto') and PASSES on golden_env (type='pct')
    try:
        if tblW is None:
            print("FAIL: Component 1 — tblW element not found in tblPr")
        else:
            actual_type = tblW.get(qn('w:type'))
            if actual_type == EXPECTED_TYPE:
                print(f"PASS: Component 1 — Table width type is 'pct' (relative percentage width) (0.5 pts)")
                total_score += 0.5
            else:
                print(f"FAIL: Component 1 — Expected width type 'pct', found '{actual_type}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table width value is 4000 (= 80% in OOXML fiftieths-of-a-percent) — 0.3 points
    # This FAILS on initial_env (w='0') and PASSES on golden_env (w='4000')
    try:
        if tblW is None:
            print("FAIL: Component 2 — tblW element not found")
        else:
            actual_w_str = tblW.get(qn('w:w'))
            if actual_w_str is not None:
                actual_w = int(actual_w_str)
                actual_pct = actual_w / 50.0
                if abs(actual_w - EXPECTED_W) <= TOLERANCE:
                    print(f"PASS: Component 2 — Table width value is {actual_w} ({actual_pct:.1f}%), expected 4000 (80.0%) (0.3 pts)")
                    total_score += 0.3
                else:
                    print(f"FAIL: Component 2 — Expected tblW w=4000 (80%), found w={actual_w} ({actual_pct:.1f}%)")
            else:
                print(f"FAIL: Component 2 — tblW w attribute not found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Cell contents are unchanged (data integrity, compound with task change) — 0.2 points
    # Checks that the 80%-width table still has correct cell data.
    # This FAILS on initial_env since Component 1 fails (not pct type), meaning
    # we gate this check on the type already having been set.
    # However, to be safe we explicitly gate it so it only scores when component 1 also passes.
    try:
        # Only award cell integrity points if the width type change was made (i.e., type is 'pct')
        # This makes this a compound condition anchored to the task change.
        if tblW is not None and tblW.get(qn('w:type')) == EXPECTED_TYPE:
            rows = table.rows
            mismatches = []
            if len(rows) != len(EXPECTED_CELLS):
                mismatches.append(f"Row count: expected {len(EXPECTED_CELLS)}, found {len(rows)}")
            else:
                for r_idx, (row, expected_row) in enumerate(zip(rows, EXPECTED_CELLS)):
                    actual_cols = [cell.text.strip() for cell in row.cells]
                    for c_idx, (actual_text, expected_text) in enumerate(zip(actual_cols, expected_row)):
                        if actual_text != expected_text:
                            mismatches.append(f"Cell[{r_idx},{c_idx}]: expected '{expected_text}', found '{actual_text}'")

            if len(mismatches) == 0:
                print(f"PASS: Component 3 — All cell contents unchanged after width change (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Cell content mismatch: {'; '.join(mismatches)}")
        else:
            # Table width type not 'pct', so this compound check cannot pass
            print(f"SKIP: Component 3 — Skipped because table width type is not 'pct' (compound condition fails)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path on VM
file_path = f'{WORKDIR}/Desktop/summary_table.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
