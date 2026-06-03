"""
Initial Setup: Table width change to 80% relative
Task ID: writer_tbl_024
Domain: libreoffice_writer

Creates a .docx file at ~/Desktop/summary_table.docx with a 3x2 table
at default/automatic width (NOT yet set to 80%).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'summary_table'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Add a short heading paragraph for context
    doc.add_heading('Summary Report', level=1)

    # Create the 3x2 table with default/automatic width
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    # Row 1 (header)
    table.cell(0, 0).text = 'Parameter'
    table.cell(0, 1).text = 'Value'

    # Row 2
    table.cell(1, 0).text = 'Temperature'
    table.cell(1, 1).text = '22\u00b0C'

    # Row 3
    table.cell(2, 0).text = 'Humidity'
    table.cell(2, 1).text = '45%'

    # Table is at default AUTO width (NOT 80% relative) — this is the pre-task state
    # Explicitly set to auto width to make sure it's not already at 80%
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
