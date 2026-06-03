"""
Reward Script: Convert straight quotes to typographic (curly) quotes
Task ID: writer_tech_073
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): All straight double quotes replaced with curly double quotes
  Component 2 (0.3): All straight single quotes/apostrophes replaced with curly equivalents
  Component 3 (0.3): Curly quotes present AND content integrity preserved (compound check)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_073'

# Unicode code points for quote characters
STRAIGHT_DOUBLE = '"'       # U+0022
STRAIGHT_SINGLE = "'"       # U+0027
LEFT_DOUBLE = '\u201c'      # "
RIGHT_DOUBLE = '\u201d'     # "
LEFT_SINGLE = '\u2018'      # '
RIGHT_SINGLE = '\u2019'     # '


def get_all_text(doc):
    """Extract all text from paragraphs and table cells."""
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return texts


def count_quotes(texts):
    """Count different quote types across all text segments."""
    combined = '\n'.join(texts)
    return {
        'straight_double': combined.count(STRAIGHT_DOUBLE),
        'straight_single': combined.count(STRAIGHT_SINGLE),
        'curly_left_double': combined.count(LEFT_DOUBLE),
        'curly_right_double': combined.count(RIGHT_DOUBLE),
        'curly_left_single': combined.count(LEFT_SINGLE),
        'curly_right_single': combined.count(RIGHT_SINGLE),
    }


def strip_quotes(text):
    """Remove all quote characters from text for content comparison."""
    for ch in [STRAIGHT_DOUBLE, STRAIGHT_SINGLE, LEFT_DOUBLE, RIGHT_DOUBLE, LEFT_SINGLE, RIGHT_SINGLE]:
        text = text.replace(ch, '')
    return text


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    texts = get_all_text(doc)
    counts = count_quotes(texts)

    print(f"Quote counts: {counts}")

    # Component 1: All straight double quotes eliminated and curly doubles present (0.4 points)
    # Initial has 110 straight double quotes and 0 curly doubles.
    # Golden should have 0 straight double and >0 curly double quotes.
    try:
        has_no_straight_double = (counts['straight_double'] == 0)
        has_curly_double = (counts['curly_left_double'] + counts['curly_right_double']) > 0

        if has_no_straight_double and has_curly_double:
            print(f"PASS: Component 1 — No straight double quotes remain, curly doubles present "
                  f"(L={counts['curly_left_double']}, R={counts['curly_right_double']}) (0.4 pts)")
            total_score += 0.4
        elif has_no_straight_double:
            # Straight doubles gone but no curly replacements (unlikely but partial credit)
            print(f"PARTIAL: Component 1 — Straight doubles removed but no curly doubles found (0.1 pts)")
            total_score += 0.1
        else:
            remaining = counts['straight_double']
            print(f"FAIL: Component 1 — {remaining} straight double quotes still present")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All straight single quotes/apostrophes eliminated and curly singles present (0.3 points)
    # Initial has 19 straight single quotes and 0 curly singles.
    # Golden should have 0 straight single and >0 curly single quotes.
    try:
        has_no_straight_single = (counts['straight_single'] == 0)
        has_curly_single = (counts['curly_left_single'] + counts['curly_right_single']) > 0

        if has_no_straight_single and has_curly_single:
            print(f"PASS: Component 2 — No straight single quotes remain, curly singles present "
                  f"(L={counts['curly_left_single']}, R={counts['curly_right_single']}) (0.3 pts)")
            total_score += 0.3
        elif has_no_straight_single:
            print(f"PARTIAL: Component 2 — Straight singles removed but no curly singles found (0.1 pts)")
            total_score += 0.1
        else:
            remaining = counts['straight_single']
            print(f"FAIL: Component 2 — {remaining} straight single quotes still present")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Quote conversion with content integrity (0.3 points)
    # COMPOUND CHECK: Curly quotes must be present (task change) AND content must be preserved.
    # This ensures we only award points when the conversion happened correctly without
    # corrupting the document. Anchored to the task change (curly quotes present).
    try:
        # First gate: curly quotes must exist (task-introduced change)
        total_curly = (counts['curly_left_double'] + counts['curly_right_double'] +
                       counts['curly_left_single'] + counts['curly_right_single'])
        curly_present = total_curly > 0

        if not curly_present:
            print(f"FAIL: Component 3 — No curly quotes found, cannot verify content integrity")
        else:
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # Check balanced curly double quotes (left count should equal right count)
            balanced_double = (counts['curly_left_double'] == counts['curly_right_double'])

            # Verify key phrases are still present (not corrupted by conversion)
            combined = '\n'.join(texts)
            key_phrases = [
                'Quarterly Software Development Report',
                'Executive Summary',
                'Horizon',
                'API Gateway Redesign',
                'Database Optimization',
                'Team Updates',
            ]
            phrases_found = sum(1 for p in key_phrases if p in combined)
            phrases_ok = (phrases_found == len(key_phrases))

            # Structure preserved
            structure_ok = (para_count == 20 and table_count == 1)

            if balanced_double and phrases_ok and structure_ok:
                print(f"PASS: Component 3 — Curly quotes balanced and content preserved "
                      f"(L_dbl={counts['curly_left_double']}, R_dbl={counts['curly_right_double']}, "
                      f"paras={para_count}, tables={table_count}, "
                      f"key_phrases={phrases_found}/{len(key_phrases)}) (0.3 pts)")
                total_score += 0.3
            elif phrases_ok:
                print(f"PARTIAL: Component 3 — Content preserved but quotes unbalanced or structure changed "
                      f"(balanced={balanced_double}, structure={structure_ok}) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 — Content corrupted during conversion "
                      f"(key_phrases={phrases_found}/{len(key_phrases)})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer — save any unsaved edits
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
