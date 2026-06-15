"""
Initial Setup: Marketing plan with messy 3-level nested list (inconsistent bullets/numbers/indentation)
Task ID: writer_mktg_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'marketing_plan_lists'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_list_paragraph(doc, text, style_name, indent_inches):
    """Add a paragraph with the given style and left indent."""
    para = doc.add_paragraph(text, style=style_name)
    para.paragraph_format.left_indent = Inches(indent_inches)
    return para


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Q3 Marketing Plan', level=0)

    # --- Introduction ---
    doc.add_paragraph(
        'This document outlines the comprehensive marketing strategy for Q3, '
        'including digital campaigns, partnership initiatives, and product launch '
        'activities across all major market segments.'
    )

    doc.add_paragraph(
        'The plan has been developed in collaboration with the sales team, '
        'product management, and senior leadership to align marketing efforts '
        'with overall business objectives for the fiscal year.'
    )

    # --- Marketing Strategy Section Heading ---
    doc.add_heading('Marketing Strategy', level=1)

    doc.add_paragraph(
        'The following outlines our key strategic priorities. '
        'Note: the nested action items require careful coordination across teams.'
    )

    # --- The INTENTIONALLY MESSY 3-level nested list ---
    # Level 1, item 1: uses a bullet
    p = doc.add_paragraph('Digital Marketing Expansion', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)  # wrong indent

    # Level 2 items under item 1 (inconsistent: dash style, wrong indent)
    p = doc.add_paragraph('Search Engine Optimization (SEO)', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.5)  # wrong - same level as parent

    # Level 3 under SEO
    p = doc.add_paragraph('Keyword research and competitive analysis', style='List Bullet 3')
    p.paragraph_format.left_indent = Inches(0.5)  # wrong - same as level 2

    p = doc.add_paragraph('On-page optimization for top 50 landing pages', style='List Number')
    p.paragraph_format.left_indent = Inches(0.75)  # mixed numbering for level 3

    p = doc.add_paragraph('Paid Search Campaigns (SEM)', style='List Number')
    p.paragraph_format.left_indent = Inches(0.25)  # wrong - same as level 1

    # Level 3 under SEM
    p = doc.add_paragraph('Budget allocation across platforms', style='List Bullet 3')
    p.paragraph_format.left_indent = Inches(1.0)

    p = doc.add_paragraph('A/B testing of ad creatives and copy', style='List Bullet 3')
    p.paragraph_format.left_indent = Inches(1.25)  # inconsistent

    p = doc.add_paragraph('Social Media Advertising', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.6)  # irregular indent

    # Level 1, item 2: uses a number
    p = doc.add_paragraph('Brand Awareness Campaigns', style='List Number')
    p.paragraph_format.left_indent = Inches(0.0)  # zero indent - wrong

    # Level 2 under Brand Awareness
    p = doc.add_paragraph('Television and Radio Spots', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.5)

    # Level 3 under TV/Radio
    p = doc.add_paragraph('Prime time slot negotiations', style='List Number 2')
    p.paragraph_format.left_indent = Inches(0.75)

    p = doc.add_paragraph('Creative production timeline and budget', style='List Bullet 3')
    p.paragraph_format.left_indent = Inches(1.0)

    p = doc.add_paragraph('Outdoor and Print Advertising', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.25)  # wrong - too shallow

    # Level 1, item 3: uses bullet again
    p = doc.add_paragraph('Partnership and Sponsorship Development', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)  # wrong - too deep

    # Level 2 under Partnerships
    p = doc.add_paragraph('Strategic Corporate Partnerships', style='List Number 2')
    p.paragraph_format.left_indent = Inches(0.75)

    # Level 3
    p = doc.add_paragraph('Joint marketing co-branding opportunities', style='List Bullet 3')
    p.paragraph_format.left_indent = Inches(0.75)  # same as parent - wrong

    p = doc.add_paragraph('Event Sponsorships', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(1.0)  # too deep

    # Level 3
    p = doc.add_paragraph('Industry conference presence and speaking slots', style='List Number')
    p.paragraph_format.left_indent = Inches(1.5)

    p = doc.add_paragraph('Community Outreach Programs', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.3)  # irregular

    # Level 1, item 4: uses bullet
    p = doc.add_paragraph('Product Launch Initiatives', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.1)  # wrong

    # Level 2 under Product Launch
    p = doc.add_paragraph('Pre-launch Buzz Building', style='List Number 2')
    p.paragraph_format.left_indent = Inches(0.6)  # irregular

    # Level 3
    p = doc.add_paragraph('Influencer seeding and early access program', style='List Bullet 3')
    p.paragraph_format.left_indent = Inches(0.6)  # same as parent - wrong

    p = doc.add_paragraph('Launch Day Execution', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.9)  # irregular

    # Level 3
    p = doc.add_paragraph('Coordinated press release and media outreach', style='List Number 2')
    p.paragraph_format.left_indent = Inches(1.2)  # irregular

    # Level 1, item 5: uses a number
    p = doc.add_paragraph('Customer Retention and Loyalty Programs', style='List Number')
    p.paragraph_format.left_indent = Inches(0.15)  # wrong - slight offset

    # Level 2 under Retention
    p = doc.add_paragraph('Loyalty Rewards System', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.4)  # too shallow

    # Level 3
    p = doc.add_paragraph('Points accumulation and redemption tiers', style='List Bullet 3')
    p.paragraph_format.left_indent = Inches(1.1)  # irregular

    p = doc.add_paragraph('Re-engagement Campaigns', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.7)  # irregular

    # Level 3
    p = doc.add_paragraph('Personalized win-back email sequences', style='List Number')
    p.paragraph_format.left_indent = Inches(0.9)  # irregular, wrong style

    p = doc.add_paragraph('Customer Success and Upsell Programs', style='List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.5)

    # --- Additional section (non-list content, should remain unchanged) ---
    doc.add_heading('Budget Overview', level=1)
    doc.add_paragraph(
        'The total marketing budget for Q3 is allocated across four main categories: '
        'digital advertising ($180,000), brand campaigns ($95,000), partnerships ($60,000), '
        'and product launches ($145,000). A contingency reserve of $20,000 is maintained '
        'for tactical opportunities.'
    )

    doc.add_heading('Timeline and Milestones', level=1)
    doc.add_paragraph(
        'Key milestones include campaign kickoff on July 1st, mid-quarter review on August 15th, '
        'and final performance assessment on September 30th. Weekly status reports will be '
        'distributed to all stakeholders every Monday morning.'
    )

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
