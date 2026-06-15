"""
Reward Script: Apply Heading 1 style to 'Installation Guide' chapter title
Task ID: writer_tech_001
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6): First paragraph has 'Heading 1' style
  Component 2 (0.4): First paragraph has 'Heading 1' AND text is 'Installation Guide'
                      AND other paragraphs were not indiscriminately changed to Heading 1
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_001'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one paragraph
    if len(doc.paragraphs) == 0:
        print("FAIL: Document has no paragraphs")
        print("REWARD: 0.0")
        return 0.0

    first_para = doc.paragraphs[0]

    # Component 1: First paragraph has 'Heading 1' style (0.6 points)
    # This is the core task requirement. In initial_env it is 'Normal'.
    try:
        style_name = first_para.style.name if first_para.style else None
        if style_name == 'Heading 1':
            print(f"PASS: Component 1 -- First paragraph style is 'Heading 1' (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 -- Expected style 'Heading 1', found '{style_name}'")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: First paragraph has 'Heading 1' AND text is 'Installation Guide'
    #              AND other paragraphs were not all changed to Heading 1 (precision check)
    # (0.4 points)
    # This compound check ensures the correct paragraph was targeted and text preserved.
    # Fails on initial_env because the style check fails.
    try:
        style_name = first_para.style.name if first_para.style else None
        text_match = first_para.text.strip() == 'Installation Guide'
        # Check that not ALL paragraphs were changed to Heading 1
        # (at least some non-Heading-1 paragraphs should remain)
        other_styles = [p.style.name for p in doc.paragraphs[1:] if p.style]
        has_non_heading1 = any(s != 'Heading 1' for s in other_styles)

        if style_name == 'Heading 1' and text_match and has_non_heading1:
            print(f"PASS: Component 2 -- Style is 'Heading 1', text is 'Installation Guide', "
                  f"other paragraphs unchanged (0.4 pts)")
            total_score += 0.4
        else:
            reasons = []
            if style_name != 'Heading 1':
                reasons.append(f"style is '{style_name}' not 'Heading 1'")
            if not text_match:
                reasons.append(f"text is '{first_para.text.strip()!r}' not 'Installation Guide'")
            if not has_non_heading1:
                reasons.append("all paragraphs changed to Heading 1 (indiscriminate)")
            print(f"FAIL: Component 2 -- {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved GUI edits
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
