"""
Initial Setup: Apply Heading 1 style to 'Installation Guide' chapter title
Task ID: writer_tech_001
Domain: libreoffice_writer

Creates a software documentation .docx with 'Installation Guide' as the first
line in Default Paragraph Style (Normal). The document contains realistic
technical documentation content below the title.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_001'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- First paragraph: 'Installation Guide' in Normal/Default style ---
    # This is the paragraph the agent needs to change to Heading 1
    para_title = doc.add_paragraph('Installation Guide')
    # Ensure it is explicitly in Normal style (default)
    para_title.style = doc.styles['Normal']

    # --- Realistic software documentation content ---
    doc.add_paragraph(
        'This document provides step-by-step instructions for installing '
        'and configuring CloudSync Pro v3.2 on your server environment. '
        'Please read the prerequisites section carefully before proceeding '
        'with the installation.'
    )

    doc.add_paragraph(
        'Prerequisites'
    )
    doc.add_paragraph(
        'Before beginning the installation, ensure your system meets the '
        'following minimum requirements:'
    )
    doc.add_paragraph('Ubuntu 22.04 LTS or later (64-bit)', style='List Bullet')
    doc.add_paragraph('8 GB RAM (16 GB recommended for production)', style='List Bullet')
    doc.add_paragraph('50 GB available disk space', style='List Bullet')
    doc.add_paragraph('Python 3.10 or later', style='List Bullet')
    doc.add_paragraph('PostgreSQL 14+ database server', style='List Bullet')
    doc.add_paragraph('Active internet connection for package downloads', style='List Bullet')

    doc.add_paragraph(
        'Step 1: Download the Installer'
    )
    doc.add_paragraph(
        'Navigate to the CloudSync Pro download portal at '
        'https://downloads.cloudsyncpro.io and select the appropriate '
        'package for your operating system. For Linux systems, download '
        'the .tar.gz archive. The current release is version 3.2.1 '
        '(build 20250315).'
    )

    doc.add_paragraph(
        'Step 2: Extract and Configure'
    )
    doc.add_paragraph(
        'Extract the downloaded archive to /opt/cloudsync using the '
        'following command:'
    )
    doc.add_paragraph(
        '    tar -xzf cloudsync-pro-3.2.1-linux-x64.tar.gz -C /opt/cloudsync'
    )
    doc.add_paragraph(
        'After extraction, run the configuration wizard by executing '
        './configure.sh from the installation directory. The wizard will '
        'prompt you for database credentials, admin email address, and '
        'the desired port number (default: 8443).'
    )

    doc.add_paragraph(
        'Step 3: Initialize the Database'
    )
    doc.add_paragraph(
        'CloudSync Pro requires a dedicated PostgreSQL database. Create '
        'the database and user with the following commands:'
    )
    doc.add_paragraph(
        '    CREATE DATABASE cloudsync_db;\n'
        '    CREATE USER cloudsync_admin WITH PASSWORD \'securepass123\';\n'
        '    GRANT ALL PRIVILEGES ON DATABASE cloudsync_db TO cloudsync_admin;'
    )

    doc.add_paragraph(
        'Troubleshooting'
    )
    doc.add_paragraph(
        'If you encounter connection errors during the database initialization '
        'step, verify that PostgreSQL is running and accepting connections on '
        'the configured port. Check the log file at /var/log/cloudsync/install.log '
        'for detailed error messages. Common issues include firewall rules blocking '
        'port 5432 and incorrect pg_hba.conf authentication settings.'
    )

    doc.add_paragraph(
        'For additional support, contact the CloudSync Pro technical team at '
        'support@cloudsyncpro.io or visit the community forum at '
        'https://community.cloudsyncpro.io.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
