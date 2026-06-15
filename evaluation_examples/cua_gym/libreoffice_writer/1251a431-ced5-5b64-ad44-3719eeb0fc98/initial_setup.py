"""
Initial Setup: Create a Writer document with a section named 'Highlights' containing three bullet points.
Task ID: writer_fs_055
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document title ---
    title = doc.add_heading("Quarterly Product Review - Q1 2025", level=1)

    # --- Introduction paragraph ---
    doc.add_paragraph(
        "This document summarizes the key findings and achievements from "
        "the first quarter product review conducted on March 28, 2025. "
        "The review covered all major product lines and their market performance."
    )

    # --- Some content before the section ---
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "During Q1 2025, the company launched three new product variants "
        "and expanded into two additional markets. Overall revenue grew by "
        "14.3% compared to the same period last year, driven primarily by "
        "the premium segment which saw a 22% increase in unit sales."
    )

    doc.add_paragraph(
        "Customer satisfaction scores remained above target at 4.6 out of "
        "5.0, with notable improvements in the onboarding experience and "
        "post-purchase support categories."
    )

    # --- Create the 'Highlights' section using a bookmark ---
    # In LibreOffice Writer, a named section is stored in .docx as
    # content between bookmark start/end markers. We create that structure.

    # Section heading
    section_heading = doc.add_heading("Highlights", level=2)

    # Add bookmark start before the bullet points
    body = doc._body._body
    # We'll wrap the bullet points in a bookmark named 'Highlights'
    bm_id = '1'

    # Insert bookmarkStart before the first bullet
    p_before = section_heading._element
    bm_start = p_before.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): bm_id,
        qn('w:name'): 'Highlights',
    })

    # Three bullet points for the Highlights section
    bullet1 = doc.add_paragraph(
        "Successfully launched the Aurora Pro line with 15,000 units "
        "sold in the first month, exceeding the initial forecast by 37%.",
        style="List Bullet"
    )
    bullet2 = doc.add_paragraph(
        "Reduced average customer support response time from 4.2 hours "
        "to 1.8 hours through the implementation of the new AI-assisted "
        "ticketing system.",
        style="List Bullet"
    )
    bullet3 = doc.add_paragraph(
        "Secured partnership agreement with Meridian Distribution Corp., "
        "opening access to 1,200 additional retail locations across the "
        "Southeast region.",
        style="List Bullet"
    )

    # Insert bookmarkEnd after the last bullet
    bm_end = p_before.makeelement(qn('w:bookmarkEnd'), {
        qn('w:id'): bm_id,
    })

    # Place bookmarkStart right before bullet1 and bookmarkEnd after bullet3
    bullet1._element.addprevious(bm_start)
    bullet3._element.addnext(bm_end)

    # --- More content after the section ---
    doc.add_heading("Next Steps", level=2)
    doc.add_paragraph(
        "The product team will focus on three priorities for Q2 2025: "
        "expanding the Aurora Pro color range, integrating the AI support "
        "system with the mobile app, and finalizing the Meridian rollout "
        "logistics. A detailed timeline will be circulated by April 15."
    )

    doc.add_paragraph(
        "Budget allocation for Q2 has been approved at $2.4 million, "
        "representing a 10% increase over Q1. The additional funding "
        "will primarily support marketing campaigns for the new retail "
        "channel and hiring two additional product engineers."
    )

    doc.add_heading("Team Acknowledgments", level=2)
    doc.add_paragraph(
        "Special thanks to the product development team led by Rachel "
        "Torres, the customer success team under David Kim, and the "
        "business development group headed by Anika Patel for their "
        "outstanding contributions this quarter."
    )

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
