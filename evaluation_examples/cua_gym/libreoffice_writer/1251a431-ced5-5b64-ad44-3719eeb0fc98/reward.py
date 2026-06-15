"""
Reward Script: Set background color of 'Highlights' section to light yellow (#FFFFF0)
Task ID: writer_fs_055
Domain: libreoffice_writer

Scoring Rubric:
  Component 1 (0.4): At least one paragraph in the Highlights section has shading with fill color FFFFF0
  Component 2 (0.3): All three bullet-point paragraphs in Highlights have the correct shading
  Component 3 (0.3): The shading is specifically FFFFF0 (light yellow) and no other paragraphs outside
                      the Highlights section have that shading applied (precision check)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_055'


def persist_app_state(domain):
    """Send Ctrl+S to save any unsaved edits in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print("PERSIST: ctrl+s sent for %s" % domain)
        except Exception as e:
            print("PERSIST_WARN: save hook failed: %s" % e)


def get_highlights_paragraph_indices(doc):
    """
    Find paragraph indices that belong to the 'Highlights' section.
    Strategy: find the paragraph with Heading style containing 'Highlights',
    then collect all subsequent paragraphs until the next heading.
    This is more robust than bookmark-based detection since bookmarks
    may be nested inside paragraphs or at body level depending on the file.
    """
    heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            if para.text.strip() == 'Highlights':
                heading_idx = i
                break

    if heading_idx is None:
        return None, "No heading paragraph with text 'Highlights' found"

    # Collect paragraphs after the heading until the next heading or end
    highlight_para_indices = []
    for j in range(heading_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[j]
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            break
        highlight_para_indices.append(j)

    if not highlight_para_indices:
        return None, "No content paragraphs found after 'Highlights' heading"

    return highlight_para_indices, None


def get_paragraph_shading_fill(para):
    """Get the w:shd fill color from paragraph properties, or None if not set."""
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    pPr = para._element.find('{%s}pPr' % ns_w)
    if pPr is None:
        return None
    shd = pPr.find('{%s}shd' % ns_w)
    if shd is None:
        return None
    fill = shd.get('{%s}fill' % ns_w)
    return fill


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Identify the Highlights section paragraphs
    highlight_indices, err = get_highlights_paragraph_indices(doc)
    if err or not highlight_indices:
        print("FAIL: Could not identify Highlights section paragraphs: %s" % err)
        print("REWARD: 0.0")
        return 0.0

    print("INFO: Highlights section paragraph indices: %s" % highlight_indices)

    # Get shading fills for highlight paragraphs
    highlight_fills = []
    for idx in highlight_indices:
        fill = get_paragraph_shading_fill(doc.paragraphs[idx])
        highlight_fills.append(fill)
        print("INFO: Para %d shading fill: %s" % (idx, fill))

    # Component 1: At least one paragraph in Highlights has FFFFF0 shading (0.4 points)
    try:
        fffff0_count = sum(1 for f in highlight_fills if f and f.upper() == 'FFFFF0')
        if fffff0_count > 0:
            print("PASS: Component 1 -- %d of %d Highlights paragraphs have FFFFF0 shading (0.4 pts)" % (fffff0_count, len(highlight_fills)))
            total_score += 0.4
        else:
            print("FAIL: Component 1 -- No Highlights paragraphs have FFFFF0 shading. Found fills: %s" % highlight_fills)
    except Exception as e:
        print("ERROR: Component 1 -- %s" % e)

    # Component 2: ALL bullet paragraphs in Highlights have FFFFF0 shading (0.3 points)
    try:
        all_shaded = all(f and f.upper() == 'FFFFF0' for f in highlight_fills)
        if all_shaded and len(highlight_fills) > 0:
            print("PASS: Component 2 -- All %d Highlights paragraphs have FFFFF0 shading (0.3 pts)" % len(highlight_fills))
            total_score += 0.3
        else:
            non_shaded = [i for i, f in zip(highlight_indices, highlight_fills) if not f or f.upper() != 'FFFFF0']
            print("FAIL: Component 2 -- Not all Highlights paragraphs shaded. Missing: para indices %s" % non_shaded)
    except Exception as e:
        print("ERROR: Component 2 -- %s" % e)

    # Component 3: No paragraphs OUTSIDE the Highlights section have FFFFF0 shading (0.3 points)
    # This ensures the shading was applied precisely to the section, not the whole document
    try:
        outside_shaded = []
        for i, para in enumerate(doc.paragraphs):
            if i not in highlight_indices:
                fill = get_paragraph_shading_fill(para)
                if fill and fill.upper() == 'FFFFF0':
                    outside_shaded.append(i)

        # This component only awards points if Component 1 passed (at least some shading exists)
        if fffff0_count > 0 and len(outside_shaded) == 0:
            print("PASS: Component 3 -- No paragraphs outside Highlights have FFFFF0 shading (0.3 pts)")
            total_score += 0.3
        elif fffff0_count == 0:
            print("FAIL: Component 3 -- No shading found at all, so precision check is moot")
        else:
            print("FAIL: Component 3 -- %d paragraphs outside Highlights also have FFFFF0 shading: indices %s" % (len(outside_shaded), outside_shaded))
    except Exception as e:
        print("ERROR: Component 3 -- %s" % e)

    final_score = min(total_score, 1.0)
    print("\nScore: %.1f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = '%s/%s.docx' % (WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
