"""
Reward Script: Apartment listing description in LibreOffice Writer
Task ID: writer_wf_065
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Title present and centered
  Component 2 (0.25): 7-field listing details table
  Component 3 (0.20): 3 property description paragraphs
  Component 4 (0.15): 8 amenity bullet items
  Component 5 (0.10): 5 nearby facilities bullet items
  Component 6 (0.10): Contact information section
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_065'


def persist_app_state(domain):
    """Try to save any open LibreOffice document before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify apartment listing document with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    all_paras = doc.paragraphs
    all_text = [p.text.strip() for p in all_paras]

    # Component 1: Title present and centered (0.20 points)
    # The title "2-Bedroom Luxury Apartment - Downtown District" must exist
    try:
        title_found = False
        for p in all_paras:
            if '2-bedroom' in p.text.lower() and 'luxury' in p.text.lower() and 'apartment' in p.text.lower() and 'downtown' in p.text.lower():
                title_found = True
                # Check centering
                alignment = p.paragraph_format.alignment
                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    print(f"PASS: Component 1 — Title found and centered: '{p.text[:60]}' (0.20 pts)")
                    total_score += 0.20
                else:
                    print(f"PARTIAL: Component 1 — Title found but not centered (alignment={alignment}) (0.10 pts)")
                    total_score += 0.10
                break
        if not title_found:
            print("FAIL: Component 1 — Title '2-Bedroom Luxury Apartment - Downtown District' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 7-field listing details table (0.25 points)
    # Table should have 7 rows with fields: Price, Size, Bedrooms, Bathrooms, Floor, Parking, Available From
    try:
        required_fields = ['price', 'size', 'bedrooms', 'bathrooms', 'floor', 'parking', 'available']
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            found_fields = 0
            for row in table.rows:
                label = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip() if len(row.cells) > 1 else ''
                for field in required_fields:
                    if field in label and len(value) > 0:
                        found_fields += 1
                        break
            if found_fields >= 7:
                print(f"PASS: Component 2 — Table with all 7 required fields found (0.25 pts)")
                total_score += 0.25
            elif found_fields >= 4:
                partial = round(0.25 * found_fields / 7, 2)
                print(f"PARTIAL: Component 2 — Table has {found_fields}/7 fields ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 2 — Table has only {found_fields}/7 required fields")
        else:
            print("FAIL: Component 2 — No table found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 3 property description paragraphs (0.20 points)
    # There should be 3 non-empty body paragraphs that form the property description
    try:
        # Find paragraphs between "Property Description" header and "Amenities" header
        desc_start = None
        desc_end = None
        for i, text in enumerate(all_text):
            if 'property description' in text.lower():
                desc_start = i + 1
            elif desc_start is not None and text.strip().lower() == 'amenities':
                desc_end = i
                break

        if desc_start is not None and desc_end is not None:
            desc_paras = [t for t in all_text[desc_start:desc_end] if len(t) > 20]
            if len(desc_paras) >= 3:
                print(f"PASS: Component 3 — 3 description paragraphs found (0.20 pts)")
                total_score += 0.20
            elif len(desc_paras) >= 1:
                partial = round(0.20 * len(desc_paras) / 3, 2)
                print(f"PARTIAL: Component 3 — {len(desc_paras)}/3 description paragraphs ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — No substantial description paragraphs found between headers")
        elif desc_start is not None:
            # No amenities header found, look for desc paragraphs after Property Description
            desc_paras = [t for t in all_text[desc_start:desc_start+10] if len(t) > 20]
            if len(desc_paras) >= 3:
                print(f"PASS: Component 3 — 3 description paragraphs found (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — Found {len(desc_paras)} description paragraphs (need 3)")
        else:
            print("FAIL: Component 3 — 'Property Description' section header not found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 8 amenity bullet items (0.15 points)
    # Bullets after "Amenities" header and before "Nearby Facilities"
    try:
        amenity_start = None
        amenity_end = None
        for i, text in enumerate(all_text):
            if text.lower() == 'amenities' or 'amenities' in text.lower() and len(text) < 30:
                amenity_start = i + 1
            elif amenity_start is not None and ('nearby' in text.lower() or 'facilities' in text.lower()):
                amenity_end = i
                break

        if amenity_start is not None:
            if amenity_end is None:
                amenity_end = len(all_text)
            # Count bullet items (non-empty paragraphs with List Bullet style or non-empty text)
            bullet_count = 0
            for j in range(amenity_start, amenity_end):
                p = all_paras[j]
                if p.text.strip() and (p.style.name in ('List Bullet', 'List Bullet 2', 'List Bullet 3') or len(p.text.strip()) > 3):
                    bullet_count += 1
            if bullet_count >= 8:
                print(f"PASS: Component 4 — {bullet_count} amenity bullet items found (0.15 pts)")
                total_score += 0.15
            elif bullet_count >= 4:
                partial = round(0.15 * bullet_count / 8, 2)
                print(f"PARTIAL: Component 4 — {bullet_count}/8 amenity items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 — Only {bullet_count} amenity items found (need 8)")
        else:
            print("FAIL: Component 4 — 'Amenities' section header not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 5 nearby facilities bullet items (0.10 points)
    try:
        facility_start = None
        facility_end = None
        for i, text in enumerate(all_text):
            if 'nearby' in text.lower() and 'facilit' in text.lower():
                facility_start = i + 1
            elif facility_start is not None and ('contact' in text.lower() or i >= facility_start + 15):
                facility_end = i
                break

        if facility_start is not None:
            if facility_end is None:
                facility_end = len(all_text)
            bullet_count = 0
            for j in range(facility_start, facility_end):
                p = all_paras[j]
                if p.text.strip() and (p.style.name in ('List Bullet', 'List Bullet 2', 'List Bullet 3') or len(p.text.strip()) > 3):
                    bullet_count += 1
            if bullet_count >= 5:
                print(f"PASS: Component 5 — {bullet_count} nearby facilities items found (0.10 pts)")
                total_score += 0.10
            elif bullet_count >= 2:
                partial = round(0.10 * bullet_count / 5, 2)
                print(f"PARTIAL: Component 5 — {bullet_count}/5 facility items ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — Only {bullet_count} facility items found (need 5)")
        else:
            print("FAIL: Component 5 — 'Nearby Facilities' section header not found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Contact information section (0.10 points)
    # Should have contact info at the bottom: agent name, phone/email
    try:
        contact_start = None
        for i, text in enumerate(all_text):
            if 'contact' in text.lower() and ('info' in text.lower() or 'detail' in text.lower()):
                contact_start = i + 1
                break

        if contact_start is not None:
            contact_paras = [t for t in all_text[contact_start:] if t]
            has_phone = any('phone' in t.lower() or '(' in t and ')' in t for t in contact_paras)
            has_email = any('email' in t.lower() or '@' in t for t in contact_paras)
            has_name = len(contact_paras) >= 1 and len(contact_paras[0]) > 3

            contact_score = 0
            if has_name:
                contact_score += 1
            if has_phone:
                contact_score += 1
            if has_email:
                contact_score += 1

            if contact_score >= 3:
                print(f"PASS: Component 6 — Contact info with name, phone, email found (0.10 pts)")
                total_score += 0.10
            elif contact_score >= 1:
                partial = round(0.10 * contact_score / 3, 2)
                print(f"PARTIAL: Component 6 — {contact_score}/3 contact elements ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — No contact information elements found")
        else:
            print("FAIL: Component 6 — 'Contact Information' section header not found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
