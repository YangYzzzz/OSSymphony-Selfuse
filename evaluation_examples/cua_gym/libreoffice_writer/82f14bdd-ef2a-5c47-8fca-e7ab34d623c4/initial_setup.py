"""
Initial Setup: Project plan document with a timeline image on page 2, no caption yet.
Task ID: writer_obj_029
Domain: libreoffice_writer

Creates /home/user/Desktop/project_plan.docx with:
  - Page 1: Project overview content
  - Page 2: Timeline diagram image (12cm x 6cm) WITHOUT a caption
"""

import os
import shlex
import subprocess
import time
import io
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_029'
OUTPUT = f'{WORKDIR}/project_plan.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_timeline_image():
    """Create a simple timeline diagram image (12cm x 6cm at 96dpi)."""
    # 12cm x 6cm at 96 dpi: 12/2.54*96 ~ 453 x 6/2.54*96 ~ 226
    width, height = 453, 226
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Draw timeline bar
    bar_y = height // 2
    bar_x_start = 30
    bar_x_end = width - 30
    draw.rectangle([bar_x_start, bar_y - 6, bar_x_end, bar_y + 6], fill=(70, 130, 180))

    # Milestones
    milestones = [
        (0.10, 'Q1\nKickoff'),
        (0.35, 'Q2\nDesign'),
        (0.60, 'Q3\nBuild'),
        (0.85, 'Q4\nLaunch'),
    ]
    for frac, label in milestones:
        x = int(bar_x_start + frac * (bar_x_end - bar_x_start))
        draw.ellipse([x - 9, bar_y - 9, x + 9, bar_y + 9], fill=(220, 80, 60), outline=(180, 40, 20), width=2)
        lines = label.split('\n')
        for i, line in enumerate(lines):
            tw = draw.textlength(line)
            draw.text((x - tw / 2, bar_y + 18 + i * 16), line, fill=(40, 40, 40))

    # Title
    title = 'Project Timeline 2025'
    tw = draw.textlength(title)
    draw.text(((width - tw) / 2, 10), title, fill=(30, 30, 30))

    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn

    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # ---- PAGE 1: Project Overview ----
    # Title
    title = doc.add_heading('Q3/Q4 Project Plan 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines the strategic project plan for the second half of 2025. '
        'It covers the key milestones, deliverables, resource allocation, and timeline '
        'for the Integrated Platform Modernization initiative across all business units.'
    )

    # Objectives
    doc.add_heading('Key Objectives', level=1)
    objectives = [
        'Migrate legacy backend systems to cloud-native microservices architecture.',
        'Deliver a redesigned customer-facing portal by end of Q3.',
        'Complete security audit and compliance certification (ISO 27001) by October.',
        'Reduce average system downtime from 2.3% to under 0.5% SLA target.',
        'Onboard three new enterprise clients to the modernized platform.',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # Project Team
    doc.add_heading('Project Team', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Role'
    hdr_cells[2].text = 'Department'
    for run in hdr_cells[0].paragraphs[0].runs:
        run.bold = True
    for run in hdr_cells[1].paragraphs[0].runs:
        run.bold = True
    for run in hdr_cells[2].paragraphs[0].runs:
        run.bold = True

    team_data = [
        ('Alexandra Torres', 'Project Manager', 'PMO'),
        ('David Kim', 'Lead Architect', 'Engineering'),
        ('Priya Nair', 'UX Lead', 'Product Design'),
        ('Samuel Okafor', 'Backend Engineer', 'Engineering'),
        ('Mei-Lin Chen', 'Security Analyst', 'InfoSec'),
        ('Jordan Westfield', 'QA Engineer', 'Quality Assurance'),
    ]
    for name, role, dept in team_data:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = role
        row[2].text = dept

    # Budget Summary
    doc.add_heading('Budget Summary', level=1)
    doc.add_paragraph(
        'The approved budget for the project is $1,240,000 for H2 2025. '
        'This covers personnel costs ($820,000), infrastructure ($280,000), '
        'third-party licensing ($95,000), and contingency reserve ($45,000).'
    )

    # ---- PAGE BREAK → PAGE 2: Timeline ----
    doc.add_page_break()

    doc.add_heading('Project Timeline', level=1)
    doc.add_paragraph(
        'The diagram below illustrates the high-level project timeline across all four '
        'quarters of 2025, including major milestones and phase transitions.'
    )

    # Insert timeline image (12cm x 6cm) — NO CAPTION
    img_bytes = create_timeline_image()
    para_img = doc.add_paragraph()
    para_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_img = para_img.add_run()
    run_img.add_picture(img_bytes, width=Cm(12), height=Cm(6))

    # Continue page 2 with more content
    doc.add_heading('Phase Breakdown', level=2)

    phases = [
        ('Phase 1 — Discovery & Planning (Jul 2025)',
         'Requirements gathering, stakeholder alignment, architecture design review, '
         'and team onboarding. Deliverables: Technical Specification Document v1.0.'),
        ('Phase 2 — Design & Prototyping (Aug 2025)',
         'UX wireframes, API contract definition, proof-of-concept builds, and '
         'security threat modelling sessions. Deliverables: Approved Design Prototype.'),
        ('Phase 3 — Development & Integration (Sep–Oct 2025)',
         'Iterative sprint development cycles, continuous integration pipeline setup, '
         'module integration testing, and performance benchmarking.'),
        ('Phase 4 — Testing & Hardening (Nov 2025)',
         'Full regression suite execution, penetration testing, user acceptance testing '
         'with pilot clients, and load testing at 3× projected peak traffic.'),
        ('Phase 5 — Deployment & Handover (Dec 2025)',
         'Phased production rollout, monitoring dashboard activation, SRE runbook '
         'documentation, and formal project closure report.'),
    ]
    for phase_title, phase_desc in phases:
        doc.add_paragraph(phase_title, style='List Bullet')
        doc.add_paragraph(phase_desc)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
