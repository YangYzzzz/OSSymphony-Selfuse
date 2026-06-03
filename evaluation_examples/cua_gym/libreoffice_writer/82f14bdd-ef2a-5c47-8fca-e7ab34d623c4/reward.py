"""
Reward Script: Add a caption below the image on page 2
Task ID: writer_obj_029
Domain: libreoffice_writer
Scoring:
  Component 1: Caption text 'Figure 1: Project Timeline Overview' exists in document (0.5 pts)
  Component 2: Caption paragraph uses 'Caption' style (0.3 pts)
  Component 3: Caption paragraph immediately follows the image paragraph (0.2 pts)
Total: 1.0
"""

import os

from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_029'
FILE_PATH = f'{WORKDIR}/project_plan.docx'

EXPECTED_CAPTION_TEXT = 'Figure 1: Project Timeline Overview'


def find_image_paragraph_index(doc):
    """Find the index of the paragraph containing an inline image (drawing)."""
    for i, para in enumerate(doc.paragraphs):
        xml = para._element.xml
        if 'w:drawing' in xml or 'pic:pic' in xml or 'a:graphicData' in xml:
            return i
    return -1


def verify_task(file_path):
    """
    Verify that a caption 'Figure 1: Project Timeline Overview' has been added
    below the image on page 2 of project_plan.docx.

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — fail fast if unreadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Component 1: Caption text exists anywhere in the document (0.5 pts) ---
    # This is the primary check: the text 'Figure 1: Project Timeline Overview'
    # should appear in a paragraph after the task is completed.
    try:
        caption_paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if EXPECTED_CAPTION_TEXT in para.text:
                caption_paragraphs.append(i)

        if caption_paragraphs:
            print(f"PASS: Component 1 — Caption text '{EXPECTED_CAPTION_TEXT}' found "
                  f"at paragraph index(es): {caption_paragraphs} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Caption text '{EXPECTED_CAPTION_TEXT}' not found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: Caption paragraph uses 'Caption' style (0.3 pts) ---
    # A properly inserted LibreOffice caption uses the built-in 'Caption' paragraph style.
    # This check requires both the text AND the correct style.
    try:
        styled_captions = []
        for i, para in enumerate(doc.paragraphs):
            if EXPECTED_CAPTION_TEXT in para.text and para.style.name == 'Caption':
                styled_captions.append(i)

        if styled_captions:
            print(f"PASS: Component 2 — Caption paragraph has 'Caption' style "
                  f"at index(es): {styled_captions} (0.3 pts)")
            total_score += 0.3
        else:
            # Check if text exists but style is wrong
            text_found = any(EXPECTED_CAPTION_TEXT in p.text for p in doc.paragraphs)
            if text_found:
                wrong_styles = [(i, p.style.name) for i, p in enumerate(doc.paragraphs)
                                if EXPECTED_CAPTION_TEXT in p.text]
                print(f"FAIL: Component 2 — Caption text found but style is not 'Caption': {wrong_styles}")
            else:
                print(f"FAIL: Component 2 — Caption text not found, cannot check style")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: Caption paragraph immediately follows the image paragraph (0.2 pts) ---
    # The caption should appear directly after the image paragraph, not elsewhere in the document.
    try:
        img_idx = find_image_paragraph_index(doc)
        if img_idx == -1:
            print("FAIL: Component 3 — Could not locate image paragraph in document")
        else:
            # Check if the paragraph immediately after the image has the caption text
            if img_idx + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[img_idx + 1]
                if EXPECTED_CAPTION_TEXT in next_para.text:
                    print(f"PASS: Component 3 — Caption immediately follows image paragraph "
                          f"(image at {img_idx}, caption at {img_idx + 1}) (0.2 pts)")
                    total_score += 0.2
                else:
                    print(f"FAIL: Component 3 — Paragraph after image is not the caption. "
                          f"Image at {img_idx}, next para text: {next_para.text!r}")
            else:
                print(f"FAIL: Component 3 — Image is the last paragraph, no following caption")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint: test against canonical artifact path on VM
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
