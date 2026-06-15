"""
Initial Setup: Employee handbook with old company name and address
Task ID: writer_hr_083
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_083'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

COMPANY = 'DataSync Solutions'
ADDRESS = '500 Tech Park Blvd, Austin, TX 78701'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, size=11, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if alignment:
        p.paragraph_format.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def setup_header_footer(section, header_text, footer_text):
    """Set header and footer for a section."""
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        hp = header.paragraphs[0]
        hp.clear()
    else:
        hp = header.add_paragraph()
    run = hp.add_run(header_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        fp = footer.paragraphs[0]
        fp.clear()
    else:
        fp = footer.add_paragraph()
    run = fp.add_run(footer_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_new_section(doc, header_text, footer_text):
    """Add a new section with its own header/footer."""
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.left_margin = Inches(1)
    new_section.right_margin = Inches(1)
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)
    setup_header_footer(new_section, header_text, footer_text)
    return new_section


def create_initial():
    doc = Document()

    # Page setup for first section
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # We will track company name count and address count.
    # Target: 35 company name instances, 8 address instances
    # Plan:
    #   Headers/Footers: 5 sections x header has company name = 5 in headers
    #                    5 sections x footer has address = 5 in footers (but only some)
    #   Body text: remaining instances

    # --- SECTION 1: COVER PAGE & TABLE OF CONTENTS ---
    # Header: company name (1)
    # Footer: address (1)
    setup_header_footer(
        section,
        f'{COMPANY} - Employee Handbook',           # company #1
        f'{COMPANY} | {ADDRESS}'                     # company #2, address #1
    )

    # Cover page - Title
    for _ in range(4):
        doc.add_paragraph()  # spacing

    # company #3
    add_para(doc, COMPANY, bold=True, size=28,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=12)
    add_para(doc, 'Employee Handbook', bold=True, size=22,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=6)
    add_para(doc, 'Fiscal Year 2025-2026', size=14,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=6)
    # company #4, address #2
    add_para(doc, f'{COMPANY} Headquarters', size=12,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=4)
    add_para(doc, ADDRESS, size=12,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=4)
    add_para(doc, 'Confidential - Internal Use Only', size=10,
             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=4)

    # Table of contents page
    doc.add_page_break()
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Welcome to DataSync Solutions ....................... 3',    # company #5
        '2. Company Overview ........................................ 4',
        '3. Employment Policies ..................................... 6',
        '4. Compensation & Benefits ............................... 8',
        '5. Workplace Conduct ....................................... 10',
        '6. Health & Safety .......................................... 12',
        '7. Information Security .................................... 14',
        '8. Performance Management ................................ 16',
        '9. Leave Policies ........................................... 18',
        '10. Termination & Separation ............................. 20',
        '11. Acknowledgment ......................................... 22',
    ]
    for item in toc_items:
        add_para(doc, item, size=11, space_after=4)

    # --- SECTION 2: WELCOME & COMPANY OVERVIEW ---
    add_new_section(doc,
                    f'{COMPANY} - Employee Handbook',           # company #6
                    f'Confidential | {COMPANY}')                # company #7

    add_heading(doc, f'1. Welcome to {COMPANY}', level=1)       # company #8
    # company #9
    add_para(doc, f'Dear New Team Member,\n\nWelcome to {COMPANY}! We are thrilled to have you join our growing team of talented professionals. This handbook has been designed to help you understand the policies, benefits, and expectations that make our workplace exceptional.')
    # company #10
    add_para(doc, f'At {COMPANY}, we believe that our employees are our greatest asset. Since our founding in 2012, we have grown from a small startup to a leading provider of enterprise data synchronization solutions, serving over 2,000 clients worldwide.')
    # company #11
    add_para(doc, f'This handbook applies to all employees of {COMPANY}, including full-time, part-time, and contract workers. Please read it carefully and reach out to Human Resources if you have any questions.')

    add_heading(doc, '2. Company Overview', level=1)
    add_heading(doc, '2.1 Mission Statement', level=2)
    # company #12
    add_para(doc, f'{COMPANY} is dedicated to empowering organizations with seamless, secure, and scalable data integration solutions. We strive to eliminate data silos and create unified ecosystems that drive business intelligence and operational excellence.')
    add_heading(doc, '2.2 Core Values', level=2)
    values = [
        'Innovation First: We constantly push the boundaries of what is possible in data technology.',
        'Customer Obsession: Every decision starts with the question, "How does this benefit our clients?"',
        'Integrity Always: We conduct business ethically and transparently in every interaction.',
        'Team Excellence: We achieve more together than any individual could alone.',
        'Continuous Learning: We invest in our people and encourage professional growth.',
    ]
    for v in values:
        doc.add_paragraph(v, style='List Bullet')

    add_heading(doc, '2.3 Office Locations', level=2)
    # company #13, address #3
    add_para(doc, f'{COMPANY} Headquarters:\n{ADDRESS}\nPhone: (512) 555-0147\nFax: (512) 555-0148')
    add_para(doc, 'Regional Offices:\n- San Francisco, CA: 200 Market Street, Suite 1500\n- New York, NY: 350 Fifth Avenue, 21st Floor\n- London, UK: 25 Old Broad Street, EC2N 1HQ')

    # --- SECTION 3: EMPLOYMENT POLICIES ---
    add_new_section(doc,
                    f'{COMPANY} - Employee Handbook',           # company #14
                    f'Confidential | {COMPANY}')                # company #15

    add_heading(doc, '3. Employment Policies', level=1)
    add_heading(doc, '3.1 Equal Employment Opportunity', level=2)
    # company #16
    add_para(doc, f'{COMPANY} is an equal opportunity employer. We do not discriminate based on race, color, religion, sex, national origin, age, disability, genetic information, veteran status, sexual orientation, gender identity, or any other characteristic protected by federal, state, or local law.')
    add_heading(doc, '3.2 Employment Classifications', level=2)
    add_para(doc, 'Full-Time Regular: Employees scheduled to work 40 hours per week on a regular basis. Eligible for all company benefits after the probationary period of 90 days.')
    add_para(doc, 'Part-Time Regular: Employees scheduled to work between 20-39 hours per week. Eligible for prorated benefits based on hours worked.')
    add_para(doc, 'Temporary/Contract: Employees hired for a specific project or defined period. Benefits eligibility varies based on contract terms.')

    add_heading(doc, '3.3 Onboarding Process', level=2)
    # company #17
    add_para(doc, f'All new hires at {COMPANY} will complete a comprehensive onboarding program during their first two weeks. This includes:')
    onboarding = [
        'Day 1: IT setup, badge issuance, workspace assignment, and welcome orientation',
        'Days 2-3: Department-specific training and introduction to team members',
        'Week 1: Compliance training (information security, workplace harassment, safety protocols)',
        'Week 2: Role-specific shadowing and initial project assignment',
    ]
    for o in onboarding:
        doc.add_paragraph(o, style='List Bullet')

    add_heading(doc, '3.4 Work Schedule & Remote Work', level=2)
    add_para(doc, 'Standard business hours are Monday through Friday, 9:00 AM to 6:00 PM local time. Flexible scheduling arrangements may be approved by department managers with HR consultation.')
    add_para(doc, 'Remote work is available for eligible positions with manager approval. Remote employees must maintain a dedicated workspace, reliable internet connection (minimum 50 Mbps), and be available during core hours (10:00 AM - 3:00 PM local time).')

    # --- SECTION 4: COMPENSATION & BENEFITS ---
    add_new_section(doc,
                    f'{COMPANY} - Employee Handbook',           # company #18
                    f'{ADDRESS} | HR Department')               # address #4

    add_heading(doc, '4. Compensation & Benefits', level=1)
    add_heading(doc, '4.1 Compensation Philosophy', level=2)
    # company #19
    add_para(doc, f'{COMPANY} is committed to providing competitive compensation that attracts, retains, and motivates top talent. Our compensation framework is reviewed annually against industry benchmarks using data from Radford, Mercer, and Glassdoor surveys.')
    add_heading(doc, '4.2 Pay Structure', level=2)
    add_para(doc, 'Salary ranges are established for each job level and are adjusted annually based on market conditions. Current pay bands:')

    # Add a salary table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Level', 'Title Range', 'Salary Range', 'Bonus Target']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    salary_data = [
        ['L1', 'Associate / Junior Engineer', '$55,000 - $75,000', '5%'],
        ['L2', 'Specialist / Engineer', '$70,000 - $95,000', '8%'],
        ['L3', 'Senior Specialist / Sr. Engineer', '$90,000 - $130,000', '12%'],
        ['L4', 'Manager / Lead Engineer', '$120,000 - $170,000', '15%'],
        ['L5', 'Director / Principal Engineer', '$160,000 - $220,000', '20%'],
    ]
    for r, row_data in enumerate(salary_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    add_para(doc, '')
    add_heading(doc, '4.3 Benefits Overview', level=2)
    # company #20
    add_para(doc, f'{COMPANY} offers a comprehensive benefits package to all eligible employees:')
    benefits = [
        'Medical Insurance: Choice of PPO or HMO plans through Blue Cross Blue Shield. Company covers 85% of employee premium and 60% of dependent premium.',
        'Dental Insurance: Delta Dental PPO plan. Company covers 80% of employee premium.',
        'Vision Insurance: VSP plan covering annual eye exams and $200 frame allowance.',
        '401(k) Retirement Plan: Company matches 100% of first 4% and 50% of next 2% of employee contributions. Immediate vesting.',
        'Life Insurance: Company-paid group term life insurance at 2x annual salary (up to $500,000).',
        'Employee Stock Purchase Plan (ESPP): Purchase company stock at 15% discount through payroll deductions.',
        'Professional Development: $3,500 annual stipend for conferences, courses, and certifications.',
        'Wellness Program: $500 annual wellness reimbursement for gym memberships, fitness classes, and health programs.',
    ]
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    # --- SECTION 5: WORKPLACE CONDUCT & HEALTH/SAFETY ---
    add_new_section(doc,
                    f'{COMPANY} - Employee Handbook',           # company #21
                    f'Confidential | {COMPANY}')                # company #22

    add_heading(doc, '5. Workplace Conduct', level=1)
    add_heading(doc, '5.1 Professional Standards', level=2)
    # company #23
    add_para(doc, f'All {COMPANY} employees are expected to conduct themselves professionally and ethically at all times. This includes interactions with colleagues, clients, vendors, and the public.')
    add_heading(doc, '5.2 Anti-Harassment Policy', level=2)
    # company #24
    add_para(doc, f'{COMPANY} is committed to providing a workplace free from harassment of any kind. Harassment includes unwelcome conduct based on protected characteristics that creates an intimidating, hostile, or offensive work environment.')
    add_para(doc, 'Any employee who experiences or witnesses harassment should report it immediately to their manager, HR, or through the anonymous ethics hotline at (512) 555-0199. All reports will be investigated promptly and confidentially.')
    add_heading(doc, '5.3 Dress Code', level=2)
    add_para(doc, 'Business casual attire is expected Monday through Thursday. Casual dress (including jeans and sneakers) is permitted on Fridays. Client-facing meetings require business professional attire.')

    add_heading(doc, '6. Health & Safety', level=1)
    add_heading(doc, '6.1 Workplace Safety', level=2)
    # company #25
    add_para(doc, f'{COMPANY} is committed to maintaining a safe and healthy workplace. All employees must comply with applicable health and safety regulations and report hazardous conditions immediately.')
    add_heading(doc, '6.2 Emergency Procedures', level=2)
    # address #5
    add_para(doc, f'Emergency evacuation routes are posted throughout the building at {ADDRESS}. Fire drills are conducted quarterly. In case of emergency, proceed to the nearest emergency exit and gather at the designated assembly point in the south parking lot.')
    add_para(doc, 'Emergency Contacts:\n- Building Security: (512) 555-0150\n- Fire/Police/Medical: 911\n- HR Emergency Line: (512) 555-0199')

    # --- SECTION 5 (cont): INFORMATION SECURITY & PERFORMANCE ---
    add_new_section(doc,
                    f'{COMPANY} - Employee Handbook',           # company #26
                    f'Confidential | {COMPANY}')                # company #27

    add_heading(doc, '7. Information Security', level=1)
    add_heading(doc, '7.1 Data Protection', level=2)
    # company #28
    add_para(doc, f'As a data technology company, {COMPANY} holds information security as a paramount responsibility. All employees must adhere to our Information Security Policy (ISP-2025) and complete annual security awareness training.')
    add_heading(doc, '7.2 Acceptable Use Policy', level=2)
    add_para(doc, 'Company-issued devices and network resources are provided for business purposes. Limited personal use is permitted provided it does not interfere with work responsibilities, compromise security, or violate any company policy.')
    add_para(doc, 'Prohibited activities include:\n- Downloading unauthorized software or applications\n- Sharing login credentials with any person, including coworkers\n- Connecting personal devices to the corporate network without IT approval\n- Storing sensitive client data on personal devices or cloud services\n- Accessing inappropriate or illegal content')

    add_heading(doc, '8. Performance Management', level=1)
    add_heading(doc, '8.1 Performance Review Cycle', level=2)
    # company #29
    add_para(doc, f'{COMPANY} conducts semi-annual performance reviews in June and December. Reviews include self-assessment, manager evaluation, and optional peer feedback through our internal platform.')
    add_heading(doc, '8.2 Rating Scale', level=2)
    perf_table = doc.add_table(rows=6, cols=3)
    perf_table.style = 'Table Grid'
    perf_headers = ['Rating', 'Description', 'Bonus Multiplier']
    for i, h in enumerate(perf_headers):
        cell = perf_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    perf_data = [
        ['5 - Exceptional', 'Consistently exceeds all expectations', '1.5x'],
        ['4 - Exceeds', 'Frequently exceeds expectations', '1.25x'],
        ['3 - Meets', 'Consistently meets all expectations', '1.0x'],
        ['2 - Developing', 'Partially meets expectations', '0.5x'],
        ['1 - Below', 'Does not meet expectations', '0x'],
    ]
    for r, row_data in enumerate(perf_data, 1):
        for c, val in enumerate(row_data):
            perf_table.cell(r, c).text = val

    add_para(doc, '')
    add_heading(doc, '8.3 Career Development', level=2)
    # company #30
    add_para(doc, f'{COMPANY} supports career growth through individual development plans (IDPs), mentorship programs, and internal mobility. Employees are encouraged to discuss career aspirations with their managers during quarterly check-ins.')
    add_para(doc, 'Internal job postings are available on the company intranet. Current employees receive priority consideration for open positions, provided they have been in their current role for at least 12 months and have a performance rating of 3 or above.')

    # --- SECTION 6: LEAVE POLICIES & TERMINATION ---
    add_new_section(doc,
                    f'{COMPANY} - Employee Handbook',           # company #31
                    f'{ADDRESS} | HR Department')               # address #6

    add_heading(doc, '9. Leave Policies', level=1)
    add_heading(doc, '9.1 Paid Time Off (PTO)', level=2)
    add_para(doc, 'PTO accrual rates are based on years of service:')
    pto_table = doc.add_table(rows=4, cols=3)
    pto_table.style = 'Table Grid'
    pto_headers = ['Years of Service', 'Annual PTO Days', 'Max Carryover']
    for i, h in enumerate(pto_headers):
        cell = pto_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    pto_data = [
        ['0 - 2 years', '15 days', '5 days'],
        ['3 - 7 years', '20 days', '8 days'],
        ['8+ years', '25 days', '10 days'],
    ]
    for r, row_data in enumerate(pto_data, 1):
        for c, val in enumerate(row_data):
            pto_table.cell(r, c).text = val

    add_para(doc, '')
    add_heading(doc, '9.2 Sick Leave', level=2)
    add_para(doc, 'All employees accrue 10 days of paid sick leave per year. Unused sick leave carries over up to a maximum balance of 30 days. Sick leave may be used for personal illness, medical appointments, or care of an immediate family member.')
    add_heading(doc, '9.3 Parental Leave', level=2)
    # company #32
    add_para(doc, f'{COMPANY} provides 16 weeks of paid parental leave for primary caregivers and 8 weeks for secondary caregivers following the birth, adoption, or foster placement of a child. Parental leave must be taken within 12 months of the qualifying event.')
    add_heading(doc, '9.4 Bereavement Leave', level=2)
    add_para(doc, 'Employees may take up to 5 days of paid bereavement leave for the death of an immediate family member (spouse, child, parent, sibling, grandparent) and up to 3 days for extended family members.')
    add_heading(doc, '9.5 Jury Duty & Voting', level=2)
    add_para(doc, 'Employees called for jury duty will receive full pay for up to 10 days. Additional unpaid leave may be granted if service extends beyond this period. Employees may take up to 2 hours of paid time off to vote on election days.')

    add_heading(doc, '10. Termination & Separation', level=1)
    add_heading(doc, '10.1 Voluntary Resignation', level=2)
    add_para(doc, 'Employees wishing to resign should provide at least two weeks written notice to their manager and HR. Managers and above are requested to provide four weeks notice. Resignation letters should be submitted via email to hr@datasyncsolutions.com.')
    add_heading(doc, '10.2 Involuntary Termination', level=2)
    # company #33
    add_para(doc, f'{COMPANY} reserves the right to terminate employment at any time with or without cause, subject to applicable law. Grounds for immediate termination include:')
    term_reasons = [
        'Gross misconduct or violation of company policy',
        'Theft, fraud, or dishonesty',
        'Violence, threats, or intimidation',
        'Willful destruction of company property',
        'Breach of confidentiality or information security policy',
    ]
    for t in term_reasons:
        doc.add_paragraph(t, style='List Bullet')

    add_heading(doc, '10.3 Exit Process', level=2)
    add_para(doc, 'Departing employees must complete the following before their last day:\n- Return all company property (laptop, badge, keys, parking pass)\n- Complete knowledge transfer documentation\n- Attend exit interview with HR\n- Sign separation agreement (if applicable)')
    # address #7
    add_para(doc, f'Final paychecks will be mailed to the employee\'s address on file or may be picked up at the HR office at {ADDRESS} within 5 business days of the separation date.')

    # --- SECTION 7: ACKNOWLEDGMENT ---
    add_new_section(doc,
                    f'{COMPANY} - Employee Handbook',           # company #34
                    f'{ADDRESS} | HR Department')               # address #8

    add_heading(doc, '11. Acknowledgment of Receipt', level=1)
    # company #35
    add_para(doc, f'I, the undersigned, acknowledge that I have received a copy of the {COMPANY} Employee Handbook dated January 2025. I understand that:')
    ack_items = [
        'This handbook outlines the policies, procedures, and benefits currently in effect.',
        'The company reserves the right to modify, revoke, or add to any policy at any time with reasonable notice.',
        'This handbook does not constitute an employment contract and does not guarantee employment for any specific duration.',
        'I am responsible for reading, understanding, and complying with the policies contained herein.',
        'I should consult with Human Resources if I have questions about any policy or procedure.',
    ]
    for a in ack_items:
        doc.add_paragraph(a, style='List Bullet')

    add_para(doc, '')
    add_para(doc, '')
    add_para(doc, '________________________________________')
    add_para(doc, 'Employee Name (Print)')
    add_para(doc, '')
    add_para(doc, '________________________________________')
    add_para(doc, 'Employee Signature')
    add_para(doc, '')
    add_para(doc, '________________________________________')
    add_para(doc, 'Date')
    add_para(doc, '')
    add_para(doc, '________________________________________')
    add_para(doc, 'HR Representative Signature')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count instances for verification
    count_company = 0
    count_address = 0
    # Body text
    for para in doc.paragraphs:
        count_company += para.text.count(COMPANY)
        count_address += para.text.count(ADDRESS)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                count_company += cell.text.count(COMPANY)
                count_address += cell.text.count(ADDRESS)
    # Headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            count_company += para.text.count(COMPANY)
            count_address += para.text.count(ADDRESS)
        for para in section.footer.paragraphs:
            count_company += para.text.count(COMPANY)
            count_address += para.text.count(ADDRESS)

    print(f'Company name instances: {count_company} (target: 35)')
    print(f'Address instances: {count_address} (target: 8)')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
