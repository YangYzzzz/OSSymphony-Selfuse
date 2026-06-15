"""
Reward Script: Replace company name and address in employee handbook
Task ID: writer_hr_083
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Company name replaced in body text (22 instances)
  Component 2 (0.20): Company name replaced in headers/footers (13 instances)
  Component 3 (0.20): Address replaced in body text (4 instances)
  Component 4 (0.15): Address replaced in headers/footers (4 instances)
  Component 5 (0.15): No residual old strings anywhere in document
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_083'

OLD_NAME = 'DataSync Solutions'
NEW_NAME = 'Quantum Dynamics Inc.'
OLD_ADDR = '500 Tech Park Blvd, Austin, TX 78701'
NEW_ADDR = '1200 Innovation Way, Denver, CO 80202'

# Expected counts from task context
EXPECTED_BODY_NAME = 22
EXPECTED_HF_NAME = 13
EXPECTED_BODY_ADDR = 4
EXPECTED_HF_ADDR = 4


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Count occurrences in body text ---
    body_old_name = 0
    body_new_name = 0
    body_old_addr = 0
    body_new_addr = 0
    for para in doc.paragraphs:
        body_old_name += para.text.count(OLD_NAME)
        body_new_name += para.text.count(NEW_NAME)
        body_old_addr += para.text.count(OLD_ADDR)
        body_new_addr += para.text.count(NEW_ADDR)

    # --- Count occurrences in headers/footers ---
    hf_old_name = 0
    hf_new_name = 0
    hf_old_addr = 0
    hf_new_addr = 0
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            for para in hf.paragraphs:
                hf_old_name += para.text.count(OLD_NAME)
                hf_new_name += para.text.count(NEW_NAME)
                hf_old_addr += para.text.count(OLD_ADDR)
                hf_new_addr += para.text.count(NEW_ADDR)

    print(f"Body: old_name={body_old_name}, new_name={body_new_name}, old_addr={body_old_addr}, new_addr={body_new_addr}")
    print(f"H/F:  old_name={hf_old_name}, new_name={hf_new_name}, old_addr={hf_old_addr}, new_addr={hf_new_addr}")

    # Component 1: Company name replaced in body text (0.30 points)
    # Requires: all 22 old names gone AND 22 new names present
    try:
        if body_old_name == 0 and body_new_name >= EXPECTED_BODY_NAME:
            print(f"PASS: Component 1 — body name replacement complete ({body_new_name} new, 0 old) (0.30 pts)")
            total_score += 0.30
        elif body_old_name == 0 and body_new_name > 0:
            # Partial: old removed but not all new present (maybe some merged/split)
            ratio = body_new_name / EXPECTED_BODY_NAME
            partial = round(0.30 * ratio, 2)
            print(f"PARTIAL: Component 1 — {body_new_name}/{EXPECTED_BODY_NAME} new names, 0 old ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — {body_old_name} old names remain, {body_new_name} new names found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Company name replaced in headers/footers (0.20 points)
    try:
        if hf_old_name == 0 and hf_new_name >= EXPECTED_HF_NAME:
            print(f"PASS: Component 2 — header/footer name replacement complete ({hf_new_name} new, 0 old) (0.20 pts)")
            total_score += 0.20
        elif hf_old_name == 0 and hf_new_name > 0:
            ratio = hf_new_name / EXPECTED_HF_NAME
            partial = round(0.20 * ratio, 2)
            print(f"PARTIAL: Component 2 — {hf_new_name}/{EXPECTED_HF_NAME} new names in H/F, 0 old ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — {hf_old_name} old names remain in H/F, {hf_new_name} new names found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Address replaced in body text (0.20 points)
    try:
        if body_old_addr == 0 and body_new_addr >= EXPECTED_BODY_ADDR:
            print(f"PASS: Component 3 — body address replacement complete ({body_new_addr} new, 0 old) (0.20 pts)")
            total_score += 0.20
        elif body_old_addr == 0 and body_new_addr > 0:
            ratio = body_new_addr / EXPECTED_BODY_ADDR
            partial = round(0.20 * ratio, 2)
            print(f"PARTIAL: Component 3 — {body_new_addr}/{EXPECTED_BODY_ADDR} new addrs, 0 old ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — {body_old_addr} old addrs remain, {body_new_addr} new addrs found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Address replaced in headers/footers (0.15 points)
    try:
        if hf_old_addr == 0 and hf_new_addr >= EXPECTED_HF_ADDR:
            print(f"PASS: Component 4 — header/footer address replacement complete ({hf_new_addr} new, 0 old) (0.15 pts)")
            total_score += 0.15
        elif hf_old_addr == 0 and hf_new_addr > 0:
            ratio = hf_new_addr / EXPECTED_HF_ADDR
            partial = round(0.15 * ratio, 2)
            print(f"PARTIAL: Component 4 — {hf_new_addr}/{EXPECTED_HF_ADDR} new addrs in H/F, 0 old ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — {hf_old_addr} old addrs remain in H/F, {hf_new_addr} new addrs found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: No residual old strings anywhere (0.15 points)
    # Comprehensive check: also look in tables if any
    try:
        total_old_name = body_old_name + hf_old_name
        total_old_addr = body_old_addr + hf_old_addr

        # Also check tables for residual old strings
        table_old_name = 0
        table_old_addr = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_old_name += cell.text.count(OLD_NAME)
                    table_old_addr += cell.text.count(OLD_ADDR)

        all_old = total_old_name + total_old_addr + table_old_name + table_old_addr
        if all_old == 0:
            print(f"PASS: Component 5 — zero residual old strings in entire document (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 5 — {all_old} residual old strings found (body+hf+tables)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist unsaved GUI edits before verification
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
