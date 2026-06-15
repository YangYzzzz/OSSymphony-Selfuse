"""
Initial Setup: Insert a footnote on the first mention of 'REST API'
Task ID: writer_tech_014
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_014'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Modern Web Services Architecture', level=0)

    # Paragraph 1 - Introduction (no mention of REST API)
    p1 = doc.add_paragraph()
    p1.add_run(
        'The evolution of distributed computing has fundamentally transformed how '
        'organizations design and deploy software systems. Over the past two decades, '
        'the shift from monolithic applications to service-oriented architectures has '
        'enabled unprecedented scalability and flexibility. Today, microservices and '
        'cloud-native platforms dominate enterprise software development, allowing teams '
        'to iterate rapidly and deploy independently.'
    )
    p1.paragraph_format.space_after = Pt(8)

    # Paragraph 2 - First mention of REST API
    p2 = doc.add_paragraph()
    p2.add_run(
        'At the core of modern web services lies the REST API, which has become the '
        'de facto standard for inter-service communication. By leveraging HTTP methods '
        'such as GET, POST, PUT, and DELETE, developers can create intuitive and '
        'stateless interfaces that are easy to consume. The widespread adoption of '
        'REST API design patterns has led to a rich ecosystem of tools, frameworks, '
        'and best practices that streamline the development process.'
    )
    p2.paragraph_format.space_after = Pt(8)

    # Paragraph 3 - Authentication
    doc.add_heading('Authentication and Authorization', level=1)
    p3 = doc.add_paragraph()
    p3.add_run(
        'Securing web services requires a robust authentication strategy. OAuth 2.0 '
        'has emerged as the industry standard for delegated authorization, enabling '
        'third-party applications to access user resources without exposing credentials. '
        'JSON Web Tokens (JWT) provide a compact, self-contained mechanism for securely '
        'transmitting claims between parties. When combined with Transport Layer Security '
        '(TLS), these protocols form a comprehensive security framework.'
    )
    p3.paragraph_format.space_after = Pt(8)

    # Paragraph 4 - Data formats
    doc.add_heading('Data Serialization Formats', level=1)
    p4 = doc.add_paragraph()
    p4.add_run(
        'JSON (JavaScript Object Notation) has largely replaced XML as the preferred '
        'data interchange format for REST API endpoints due to its lightweight syntax '
        'and native compatibility with JavaScript. However, Protocol Buffers (protobuf) '
        'and MessagePack offer superior performance for high-throughput systems where '
        'bandwidth and parsing speed are critical concerns. GraphQL has also gained '
        'traction as an alternative query language that allows clients to request exactly '
        'the data they need.'
    )
    p4.paragraph_format.space_after = Pt(8)

    # Paragraph 5 - Best practices
    doc.add_heading('API Design Best Practices', level=1)
    p5 = doc.add_paragraph()
    p5.add_run(
        'Well-designed APIs follow consistent naming conventions, use appropriate HTTP '
        'status codes, and provide comprehensive documentation. Versioning strategies '
        'such as URI path versioning (e.g., /v1/resources) or header-based versioning '
        'ensure backward compatibility as APIs evolve. Rate limiting, pagination, and '
        'caching mechanisms are essential for maintaining performance under load. '
        'Additionally, adopting the OpenAPI Specification (formerly Swagger) enables '
        'automated code generation and interactive documentation.'
    )
    p5.paragraph_format.space_after = Pt(8)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
