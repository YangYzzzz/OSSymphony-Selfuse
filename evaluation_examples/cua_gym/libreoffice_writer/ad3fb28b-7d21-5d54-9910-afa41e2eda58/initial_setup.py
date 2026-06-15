"""
Initial Setup: Legal letter document with no headers configured
Task ID: writer_legal_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard US Letter with 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # No headers configured - this is the initial state requirement

    # --- PAGE 1: Firm letterhead and opening ---
    # Firm name (just in body, not header)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("MITCHELL & ASSOCIATES, LLP")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Attorneys at Law")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("1250 Connecticut Avenue NW, Suite 800\nWashington, DC 20036")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Tel: (202) 555-4200 | Fax: (202) 555-4201")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"

    # Separator line
    doc.add_paragraph("_" * 72)

    # Date and addressee
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("March 28, 2026")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    lines = [
        "Robert K. Harrison, Esq.",
        "Harrison, Blake & Partners",
        "2000 K Street NW, Suite 400",
        "Washington, DC 20006",
    ]
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Re: Smith v. Jones Corp., Case No. 2025-CV-04817")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Dear Mr. Harrison:")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    # --- Letter body paragraphs (enough content to span 4 pages) ---
    body_paragraphs = [
        "I am writing on behalf of my client, Patricia Smith, in connection with the "
        "above-referenced matter. As you are aware, this firm represents Ms. Smith in her "
        "claims against Jones Corporation arising from the wrongful termination of her "
        "employment on September 15, 2025, and the subsequent breach of the severance "
        "agreement dated January 3, 2024.",

        "Following our initial review of the documents produced in the first round of "
        "discovery, we have identified several significant deficiencies in your client's "
        "production that must be addressed before we can proceed to depositions. "
        "Specifically, we note the following categories of documents that appear to be "
        "missing from the production:",

        "First, the employment records for the period from January 2023 through September "
        "2025 are incomplete. Your client's production included only select performance "
        "reviews and failed to include the complete personnel file as requested in our "
        "First Set of Interrogatories, specifically Interrogatory Nos. 4 through 12. The "
        "personnel file should include, at minimum, all written evaluations, disciplinary "
        "notices, commendation letters, training records, and internal memoranda concerning "
        "Ms. Smith's employment status.",

        "Second, we have not received any communications between the senior management "
        "team—specifically, CEO Margaret Jones, COO David Chen, and VP of Human Resources "
        "Sandra Williams—regarding the decision to terminate Ms. Smith's position. Our "
        "Request for Production No. 7 specifically sought all internal emails, memoranda, "
        "text messages, and other electronic communications regarding the restructuring "
        "of the Northeast Regional Division, which was cited as the basis for the "
        "termination. These documents are central to our client's claim that the stated "
        "reason for the termination was pretextual.",

        "Third, the financial records produced to date do not include the quarterly budget "
        "reports for the Northeast Regional Division for fiscal years 2024 and 2025. These "
        "records are essential to evaluating Jones Corporation's claim that the elimination "
        "of Ms. Smith's position was driven by financial necessity rather than retaliatory "
        "animus. Our forensic accounting expert, Dr. Richard Yamamoto of Yamamoto Financial "
        "Consulting, has reviewed the partial financial disclosures and identified several "
        "inconsistencies that warrant further investigation.",

        "Fourth, your client has not produced the complete minutes from the Board of "
        "Directors meetings held on June 12, 2025, August 3, 2025, and September 10, 2025. "
        "According to publicly available SEC filings, these meetings addressed workforce "
        "restructuring initiatives. The minutes are directly relevant to establishing the "
        "timeline of events leading to Ms. Smith's termination and the corporate "
        "decision-making process.",

        "We also note that your client's privilege log, produced on March 1, 2026, contains "
        "over 200 entries that appear to claim attorney-client privilege or work product "
        "protection for documents that, based on their descriptions, may not be properly "
        "subject to such claims. For example, Entry Nos. 47 through 63 describe general "
        "business communications between non-attorney employees regarding operational "
        "matters. Simply copying in-house counsel on routine business emails does not "
        "convert those communications into privileged material under the applicable standard "
        "set forth in In re Grand Jury Subpoena, 357 F.3d 900 (9th Cir. 2004).",

        "Furthermore, we wish to bring to your attention the deposition schedule for the "
        "upcoming months. We propose to depose the following witnesses in the order listed "
        "below, and we request that you confirm their availability by April 15, 2026:",

        "1. Margaret Jones, CEO of Jones Corporation — regarding the corporate restructuring "
        "decision and her direct involvement in the termination decision.\n"
        "2. Sandra Williams, VP of Human Resources — regarding the HR process followed "
        "in connection with Ms. Smith's termination and the adequacy of the severance offer.\n"
        "3. David Chen, COO — regarding the financial justification for the position "
        "elimination and the subsequent hiring of a replacement employee.\n"
        "4. Thomas Rivera, Regional Director — regarding Ms. Smith's day-to-day performance "
        "and his communications with senior management about the Northeast Division.\n"
        "5. Karen Mitchell, Assistant HR Director — regarding the internal investigation "
        "conducted after Ms. Smith filed her EEOC complaint.",

        "With respect to the pending motion for protective order filed by your client on "
        "March 15, 2026, we intend to file our opposition brief by April 5, 2026. We "
        "believe the motion is without merit, as the discovery requests at issue are "
        "narrowly tailored to the claims and defenses in this action and do not seek "
        "confidential trade secrets or proprietary business information. The documents "
        "requested relate solely to employment decisions and internal communications "
        "about workforce management, which are squarely within the scope of discovery "
        "under Federal Rule of Civil Procedure 26(b)(1).",

        "Additionally, we would like to discuss the possibility of scheduling a mediation "
        "session before the discovery cutoff date of July 31, 2026. While our client "
        "remains committed to pursuing her claims through trial if necessary, she is open "
        "to a good-faith effort to resolve this matter through alternative dispute "
        "resolution. We would suggest engaging the Honorable Patricia Watkins (Ret.) of "
        "JAMS, who has extensive experience mediating employment discrimination cases in "
        "this jurisdiction.",

        "In the interest of moving this case forward efficiently, we propose a meet-and-confer "
        "conference for the week of April 7, 2026, to discuss the discovery deficiencies "
        "outlined above, the deposition schedule, and the mediation proposal. Please advise "
        "as to your client's availability for such a conference.",

        "Finally, please be advised that if the outstanding discovery deficiencies are not "
        "cured within twenty-one (21) days of the date of this letter, we will have no "
        "alternative but to file a motion to compel pursuant to Federal Rule of Civil "
        "Procedure 37(a), with a request for sanctions including reasonable attorney's fees "
        "and costs incurred in connection with the motion. We trust that this step will not "
        "be necessary and that your client will comply with its discovery obligations in a "
        "timely manner.",

        "We look forward to your prompt response to the matters raised in this letter. "
        "Should you have any questions or wish to discuss any of these issues, please do "
        "not hesitate to contact me directly at (202) 555-4215 or via email at "
        "j.mitchell@mitchellassociates.com.",
    ]

    for text in body_paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15

    # Closing
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Sincerely,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("James T. Mitchell")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Senior Partner")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Mitchell & Associates, LLP")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("cc: Patricia Smith")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True

    p = doc.add_paragraph()
    run = p.add_run("     Court File")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
