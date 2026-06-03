"""
Reward Script: Verify different first page header setup in a legal document
Task ID: writer_legal_031
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): different_first_page_header_footer is enabled
  Component 2 (0.30): First page header contains 'MITCHELL & ASSOCIATES, LLP' centered
  Component 3 (0.25): Default header contains 'Smith v. Jones Corp.' (left-aligned)
  Component 4 (0.25): Default header has a PAGE field code (right-aligned via tab)
"""

import os

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_031'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) < 1:
        print("CRITICAL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: different_first_page_header_footer is enabled (0.20 points)
    # This is the key setting that allows separate first-page vs subsequent-page headers.
    # In the initial file this is False; in the golden file it must be True.
    try:
        diff_first = section.different_first_page_header_footer
        if diff_first:
            print(f"PASS: Component 1 — different_first_page_header_footer is True (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — different_first_page_header_footer is {diff_first}, expected True")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: First page header contains 'MITCHELL & ASSOCIATES, LLP' centered (0.30 points)
    # The task requires first page header showing the firm letterhead centered.
    try:
        first_hdr = section.first_page_header
        first_hdr_text = ""
        first_hdr_centered = False

        if first_hdr and first_hdr.paragraphs:
            # Collect all paragraph text from first page header
            first_hdr_text = " ".join(p.text.strip() for p in first_hdr.paragraphs if p.text.strip())

            # Check alignment — at least one paragraph with the firm name should be centered
            for p in first_hdr.paragraphs:
                if 'MITCHELL' in p.text.upper() and 'ASSOCIATES' in p.text.upper():
                    alignment = p.paragraph_format.alignment
                    # Check via python-docx API
                    if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        first_hdr_centered = True
                    # Also check raw XML for jc=center
                    pPr = p._element.find(qn('w:pPr'))
                    if pPr is not None:
                        jc = pPr.find(qn('w:jc'))
                        if jc is not None and jc.get(qn('w:val')) == 'center':
                            first_hdr_centered = True

        has_firm_name = 'MITCHELL' in first_hdr_text.upper() and 'ASSOCIATES' in first_hdr_text.upper() and 'LLP' in first_hdr_text.upper()

        if has_firm_name and first_hdr_centered:
            print(f"PASS: Component 2 — First page header has 'MITCHELL & ASSOCIATES, LLP' centered (0.30 pts)")
            total_score += 0.30
        elif has_firm_name:
            print(f"PARTIAL: Component 2 — Firm name found but not centered (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — First page header text: '{first_hdr_text}', expected 'MITCHELL & ASSOCIATES, LLP' centered")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Default header contains 'Smith v. Jones Corp.' (0.25 points)
    # Subsequent pages should show the case name left-aligned in the header.
    try:
        hdr = section.header
        hdr_text = ""
        has_case_name = False

        if hdr and hdr.paragraphs:
            hdr_text = " ".join(p.text.strip() for p in hdr.paragraphs if p.text.strip())
            # Check for the case name in the default header
            if 'Smith v. Jones Corp.' in hdr_text or 'smith v. jones corp.' in hdr_text.lower():
                has_case_name = True

        if has_case_name:
            print(f"PASS: Component 3 — Default header contains 'Smith v. Jones Corp.' (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — Default header text: '{hdr_text}', expected 'Smith v. Jones Corp.'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Default header has a PAGE field code (right-aligned via tab) (0.25 points)
    # The task requires a page number right-aligned in the default header.
    # We check for w:fldChar + w:instrText containing PAGE in the default header XML.
    try:
        hdr = section.header
        hdr_xml = hdr._element.xml

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Check for PAGE field code in header
        has_page_field = False
        for instr in hdr._element.findall('.//w:instrText', ns):
            if instr.text and 'PAGE' in instr.text.upper():
                has_page_field = True
                break

        # Check for right tab stop in header paragraph (indicating right-alignment of page number)
        has_right_tab = False
        for tab in hdr._element.findall('.//w:tab', ns):
            val = tab.get(qn('w:val'))
            if val == 'right':
                has_right_tab = True
                break

        if has_page_field and has_right_tab:
            print(f"PASS: Component 4 — Default header has PAGE field with right tab stop (0.25 pts)")
            total_score += 0.25
        elif has_page_field:
            print(f"PARTIAL: Component 4 — PAGE field found but no right tab stop (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 4 — No PAGE field code found in default header")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification (LibreOffice may have unsaved changes)
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
