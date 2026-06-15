"""
Initial Setup: Insert custom document properties and footer field
Task ID: writer_struct_077
Domain: libreoffice_writer

Creates the initial state of ml_project_doc.docx:
- A 6-page ML project document with realistic content
- NO custom document properties
- Footer is enabled but empty
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_077'
OUTPUT = f'{WORKDIR}/ml_project_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    return para


def add_paragraph(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def create_initial():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Enable footer but leave it empty (no text, no fields)
    footer = section.footer
    footer.is_linked_to_previous = False
    # Footer paragraph exists but is empty - ensure it's truly empty
    if footer.paragraphs:
        for para in footer.paragraphs:
            # Clear any existing content
            for run in para.runs:
                run.text = ''

    # ===================== PAGE 1: Title and Executive Summary =====================
    title_para = doc.add_heading('Machine Learning Project Documentation', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle_para = doc.add_paragraph('Advanced Neural Architecture for Predictive Analytics')
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.runs[0].bold = True
    subtitle_para.runs[0].font.size = Pt(14)

    meta_para = doc.add_paragraph('Version 2.3 | March 2025 | Internal Use Only')
    meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    add_heading(doc, '1. Executive Summary', level=1)
    add_paragraph(doc,
        'This document outlines the development, training, and evaluation of a next-generation '
        'machine learning model designed to enhance predictive analytics capabilities across '
        'enterprise deployments. The project commenced in Q3 2024 and targets production readiness '
        'by Q2 2025, with an estimated accuracy improvement of 23.7% over baseline models.')

    add_paragraph(doc,
        'The proposed architecture leverages transformer-based attention mechanisms combined with '
        'custom convolutional layers to process heterogeneous data streams. Initial benchmarking '
        'on the CIFAR-100 and ImageNet-1K datasets demonstrates state-of-the-art performance with '
        'significantly reduced inference latency (avg. 12ms per prediction on A100 GPUs).')

    add_paragraph(doc,
        'Key deliverables include: (1) a fully trained model checkpoint with 89.4% validation '
        'accuracy, (2) an optimized ONNX export suitable for edge deployment, and (3) comprehensive '
        'documentation for integration teams. The model was validated across 14 industry-standard '
        'benchmarks with consistent performance above the 95th percentile.')

    doc.add_page_break()

    # ===================== PAGE 2: Project Overview =====================
    add_heading(doc, '2. Project Overview', level=1)

    add_heading(doc, '2.1 Background and Motivation', level=2)
    add_paragraph(doc,
        'The data science team at TechNova Analytics identified a critical gap in real-time '
        'predictive modeling for supply chain disruption events. Traditional statistical approaches '
        'achieved only 67% precision on historical disruption data, falling short of the 85% '
        'threshold required for automated decision-making in procurement workflows.')

    add_paragraph(doc,
        'Following a comprehensive literature review of recent advances in sequence modeling '
        'and multi-modal learning, the team proposed a hybrid architecture combining LSTM-based '
        'temporal encoders with a cross-attention fusion layer. This approach was validated '
        'conceptually using three proof-of-concept experiments conducted between August and '
        'October 2024.')

    add_heading(doc, '2.2 Objectives', level=2)
    objectives = [
        'Develop a production-grade ML pipeline capable of processing 50,000 events per second',
        'Achieve at least 87% precision and 84% recall on the internal validation dataset',
        'Reduce model inference latency to under 15 milliseconds for 99th percentile requests',
        'Implement automated retraining triggers based on concept drift detection',
        'Establish monitoring dashboards with real-time performance tracking',
        'Document all model cards, bias assessments, and deployment runbooks',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    add_heading(doc, '2.3 Scope and Constraints', level=2)
    add_paragraph(doc,
        'The project scope encompasses data ingestion from five primary sources: transactional '
        'ERP systems, external logistics APIs, weather data feeds, geopolitical risk indices, '
        'and commodity price streams. Out of scope are natural language processing components '
        'for unstructured news analysis, which are planned for Phase 2 (Q4 2025).')

    add_paragraph(doc,
        'Key constraints include GDPR compliance requirements for EU customer data, a maximum '
        'model size of 2.1GB for on-premises deployment, and mandatory compatibility with '
        'Python 3.10+ and CUDA 11.8+. All experiments were conducted on a dedicated GPU cluster '
        'consisting of 8x NVIDIA A100 80GB nodes.')

    doc.add_page_break()

    # ===================== PAGE 3: Methodology =====================
    add_heading(doc, '3. Methodology', level=1)

    add_heading(doc, '3.1 Data Collection and Preprocessing', level=2)
    add_paragraph(doc,
        'Raw data was collected from production systems spanning a 36-month period (January 2022 '
        'through December 2024), yielding a dataset of approximately 847 million events. Data '
        'quality assessment revealed 3.2% missing values and 0.8% duplicate records, which were '
        'addressed through a multi-stage cleaning pipeline using custom Spark jobs.')

    # Table: Data Sources
    add_paragraph(doc, 'Table 1: Data Source Summary')
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers_row = table.rows[0]
    for cell, text in zip(headers_row.cells, ['Data Source', 'Record Count', 'Coverage Period', 'Quality Score']):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    data_rows = [
        ['ERP Transactions', '312,450,000', '2022-01 to 2024-12', '98.7%'],
        ['Logistics API', '189,230,000', '2022-03 to 2024-12', '97.2%'],
        ['Weather Feeds', '214,890,000', '2022-01 to 2024-12', '99.1%'],
        ['Risk Indices', '87,340,000', '2022-06 to 2024-12', '95.4%'],
        ['Commodity Prices', '43,090,000', '2022-01 to 2024-12', '98.9%'],
    ]
    for i, row_data in enumerate(data_rows, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    doc.add_paragraph()

    add_heading(doc, '3.2 Feature Engineering', level=2)
    add_paragraph(doc,
        'A total of 247 features were engineered across six categories: temporal lag features '
        '(72 features), rolling statistical aggregates (48 features), cross-source interaction '
        'terms (61 features), cyclical encodings for seasonal patterns (18 features), categorical '
        'embeddings (31 features), and external signal transforms (17 features).')

    add_paragraph(doc,
        'Feature selection was performed using a combination of permutation importance scores '
        'from gradient-boosted baseline models and SHAP value analysis. The final feature set '
        'of 183 features was validated through 5-fold cross-validation, achieving a mean AUC-ROC '
        'of 0.924 with standard deviation of 0.008.')

    add_heading(doc, '3.3 Model Architecture', level=2)
    add_paragraph(doc,
        'The core architecture consists of three main components: (1) a Temporal Encoder module '
        'using bidirectional LSTM layers with 512 hidden units and variational dropout (p=0.15), '
        '(2) a Multi-Head Cross-Attention module with 8 attention heads operating on a 256-dimensional '
        'key-query-value space, and (3) a Prediction Head comprising three fully-connected layers '
        'with GELU activations and layer normalization.')

    doc.add_page_break()

    # ===================== PAGE 4: Training and Evaluation =====================
    add_heading(doc, '4. Training and Evaluation', level=1)

    add_heading(doc, '4.1 Training Configuration', level=2)
    add_paragraph(doc,
        'Model training was conducted using AdamW optimizer with an initial learning rate of '
        '2e-4, cosine annealing schedule with warm restarts (T_0=10, T_mult=2), and gradient '
        'clipping at norm 1.0. The batch size was set to 512 samples per GPU, with gradient '
        'accumulation over 4 steps yielding an effective batch size of 16,384.')

    add_paragraph(doc,
        'Training proceeded for 120 epochs with early stopping patience of 15 epochs. Mixed '
        'precision training (FP16) was enabled to reduce memory footprint and accelerate '
        'computation. The final model checkpoint corresponds to epoch 97, where validation '
        'loss reached its minimum of 0.1847.')

    add_heading(doc, '4.2 Evaluation Results', level=2)
    add_paragraph(doc,
        'The trained model was evaluated on a held-out test set of 2.3 million samples not '
        'seen during training or validation. Key performance metrics are summarized in Table 2.')

    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Table Grid'
    for cell, text in zip(table2.rows[0].cells, ['Metric', 'Value', 'Benchmark']):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    metrics = [
        ['Precision', '89.3%', '85.0% target'],
        ['Recall', '86.1%', '84.0% target'],
        ['F1 Score', '87.7%', '84.5% target'],
        ['AUC-ROC', '0.9412', '0.900 target'],
        ['Inference Latency (p99)', '11.8ms', '15.0ms target'],
        ['Model Size (ONNX)', '1.87GB', '2.1GB limit'],
    ]
    for i, row_data in enumerate(metrics, 1):
        for j, val in enumerate(row_data):
            table2.rows[i].cells[j].text = val

    doc.add_paragraph()

    add_heading(doc, '4.3 Ablation Studies', level=2)
    add_paragraph(doc,
        'Ablation experiments were conducted to quantify the contribution of each architectural '
        'component. Removing the cross-attention module reduced F1 score by 4.2 points (87.7% → 83.5%). '
        'Replacing bidirectional LSTM with unidirectional reduced recall by 3.1 points. '
        'Disabling variational dropout increased overfitting, with train/val gap growing from 1.8% to 6.4%.')

    doc.add_page_break()

    # ===================== PAGE 5: Deployment and Integration =====================
    add_heading(doc, '5. Deployment and Integration', level=1)

    add_heading(doc, '5.1 Infrastructure Requirements', level=2)
    add_paragraph(doc,
        'Production deployment targets a Kubernetes cluster with autoscaling enabled. Minimum '
        'resource requirements per pod: 4 vCPU, 16GB RAM, and an NVIDIA T4 GPU (16GB VRAM) for '
        'accelerated inference. The inference service is containerized using Docker (base image: '
        'nvcr.io/nvidia/pytorch:24.03-py3) and exposed via a gRPC endpoint.')

    add_paragraph(doc,
        'A model registry entry has been created in MLflow (run ID: ml-2025-042-prod-v2.3) '
        'with all artifacts, parameters, and metrics logged. The ONNX export was validated '
        'against the PyTorch reference implementation with a maximum absolute difference of '
        '1.2e-6 across 10,000 test inputs.')

    add_heading(doc, '5.2 API Specification', level=2)
    add_paragraph(doc,
        'The prediction service exposes a RESTful API at /api/v2/predict accepting JSON payloads '
        'with a maximum size of 1MB. Authentication is handled via JWT tokens issued by the '
        'enterprise identity provider. Rate limiting is enforced at 500 requests per second '
        'per API key, with burst capacity of 2,000 requests.')

    add_paragraph(doc,
        'Input validation enforces strict schema compliance: all 183 feature fields must be '
        'present; missing values trigger a 422 error response with field-level error messages. '
        'Prediction responses include the binary class label, probability score, confidence '
        'interval (95%), and top-5 feature contributions ranked by SHAP magnitude.')

    add_heading(doc, '5.3 Monitoring and Alerting', level=2)
    add_paragraph(doc,
        'Model performance monitoring is implemented using a custom Grafana dashboard with '
        'three alert tiers: WARNING (precision drops below 87%), CRITICAL (precision drops below '
        '83%), and EMERGENCY (service availability below 99.5%). Concept drift is detected '
        'using Population Stability Index (PSI) with a threshold of 0.2.')

    add_paragraph(doc,
        'Automated retraining is triggered when any of the following conditions are met: '
        'PSI exceeds 0.25 on more than 30% of features, F1 score on shadow evaluation falls '
        'below 85%, or 60 days have elapsed since the last training run. Retraining pipelines '
        'are orchestrated via Apache Airflow DAG ml-retrain-prod-v2.')

    doc.add_page_break()

    # ===================== PAGE 6: Conclusions and Future Work =====================
    add_heading(doc, '6. Conclusions and Future Work', level=1)

    add_heading(doc, '6.1 Summary', level=2)
    add_paragraph(doc,
        'This project successfully delivered a production-grade machine learning system that '
        'exceeds all performance targets established in the project charter. The hybrid '
        'LSTM-Attention architecture demonstrates robust generalization across diverse data '
        'distributions and maintains stable performance under concept drift conditions tested '
        'in simulation.')

    add_paragraph(doc,
        'The end-to-end ML pipeline, from data ingestion through model serving, has been '
        'fully automated and integrated with existing enterprise infrastructure. Total project '
        'duration was 7 months (August 2024 through March 2025), with a team of 6 data '
        'scientists, 3 ML engineers, and 2 platform engineers.')

    add_heading(doc, '6.2 Future Work', level=2)
    future_items = [
        'Phase 2 NLP module for unstructured news feed analysis (planned Q4 2025)',
        'Federated learning implementation to enable cross-tenant model improvement',
        'Quantization experiments targeting INT8 deployment for edge devices',
        'Active learning framework to reduce labeling costs for rare disruption events',
        'Multi-task learning extension to handle four additional prediction targets',
    ]
    for item in future_items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '6.3 Acknowledgments', level=2)
    add_paragraph(doc,
        'The project team acknowledges contributions from the Data Infrastructure group '
        '(Dr. Jennifer Matsuda, team lead), Cloud Platform Engineering (Rajesh Krishnamurthy), '
        'and Security & Compliance (Lisa Okonkwo). External partnerships with DataBridge Inc. '
        'and CloudML Solutions were instrumental in accelerating the data pipeline development.')

    add_paragraph(doc,
        'This work was partially supported by an internal innovation grant from TechNova '
        'Analytics R&D budget (Grant reference: TNRD-2024-ML-007). All code, data, and '
        'model artifacts are proprietary to TechNova Analytics and subject to IP protection.')

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
