"""
Reward Script: Two-section document with different column layouts and background colors
Task ID: writer_rd_068
Domain: libreoffice_writer
Scoring:
  Component 1: Document has 2 sections (0.15 pts)
  Component 2: Section break is continuous type (0.10 pts)
  Component 3: First section has single-column layout (0.10 pts)
  Component 4: Second section has two-column layout with separator line (0.25 pts)
  Component 5: First section paragraphs have light yellow (#FFFFF0) background (0.20 pts)
  Component 6: Second section paragraphs have light blue (#F0F8FF) background (0.20 pts)
"""

import os
from docx import Document
from docx.oxml.ns import qn
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_068'


def color_distance_hex(hex1, hex2):
    """Compute Euclidean distance between two hex color strings (e.g. 'FFFFF0')."""
    r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
    r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def get_para_shading_fill(para):
    """Get the fill color hex string from a paragraph's shading, or None."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    return fill


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)
    num_paras = len(doc.paragraphs)

    # Component 1: Document has exactly 2 sections (0.15 points)
    # Initial has 1 section; golden has 2 sections.
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 — Document has {num_sections} sections (>=2) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 1 — Expected >= 2 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Section break is continuous type (0.10 points)
    # The section break separating the two sections should be continuous, not new page.
    try:
        if num_sections >= 2:
            # In python-docx, sections[0] is the first section.
            # The section break is defined in the sectPr of the first section (embedded in paragraph pPr).
            # Check start_type of sections; both should be continuous.
            sec0_type = doc.sections[0].start_type
            sec1_type = doc.sections[1].start_type
            # WD_SECTION_START.CONTINUOUS = 0
            if sec0_type is not None and sec0_type == 0:
                print(f"PASS: Component 2 — Section break is continuous (type={sec0_type}) (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 2 — Section 0 start_type={sec0_type}, expected CONTINUOUS (0)")
        else:
            print(f"FAIL: Component 2 — Cannot check section break type with only {num_sections} section(s)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: First section has single-column layout (0.10 points)
    # Initial has no explicit column setting; golden section 0 has cols num=1.
    # We check that the first section does NOT have num > 1.
    try:
        if num_sections >= 2:
            sect_pr = doc.sections[0]._sectPr
            cols_elem = sect_pr.find(qn('w:cols'))
            if cols_elem is not None:
                num_cols = cols_elem.get(qn('w:num'))
                if num_cols is None or int(num_cols) == 1:
                    print(f"PASS: Component 3 — First section has single column (num={num_cols}) (0.10 pts)")
                    total_score += 0.10
                else:
                    print(f"FAIL: Component 3 — First section has {num_cols} columns, expected 1")
            else:
                # No cols element means default single column
                print(f"PASS: Component 3 — First section has default single column (0.10 pts)")
                total_score += 0.10
        else:
            print(f"FAIL: Component 3 — Cannot check column layout with only {num_sections} section(s)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Second section has two-column layout with column separator (0.25 points)
    # Initial has no second section at all. Golden has cols num=2 and sep=1.
    try:
        if num_sections >= 2:
            sect_pr = doc.sections[1]._sectPr
            cols_elem = sect_pr.find(qn('w:cols'))
            comp4_score = 0.0
            if cols_elem is not None:
                num_cols_str = cols_elem.get(qn('w:num'))
                sep_str = cols_elem.get(qn('w:sep'))
                num_cols = int(num_cols_str) if num_cols_str else 1
                sep_val = int(sep_str) if sep_str else 0

                if num_cols == 2:
                    comp4_score += 0.15
                    print(f"  SUB-PASS: Second section has 2 columns")
                else:
                    print(f"  SUB-FAIL: Second section has {num_cols} columns, expected 2")

                if sep_val == 1:
                    comp4_score += 0.10
                    print(f"  SUB-PASS: Column separator line is enabled (sep=1)")
                else:
                    print(f"  SUB-FAIL: Column separator is {sep_val}, expected 1")
            else:
                print(f"  SUB-FAIL: No cols element in second section")

            if comp4_score > 0:
                print(f"PASS: Component 4 — Two-column layout verified ({comp4_score} pts)")
                total_score += comp4_score
            else:
                print(f"FAIL: Component 4 — Second section does not have two-column layout")
        else:
            print(f"FAIL: Component 4 — Cannot check second section with only {num_sections} section(s)")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: First section paragraphs have light yellow (#FFFFF0) background (0.20 points)
    # Initial has no paragraph shading. Golden has FFFFF0 on paras 0-3.
    # We check that paragraphs in the first section (before the section break) have the right fill.
    try:
        if num_sections >= 2:
            # Find which paragraph contains the section break (sectPr in pPr)
            break_para_idx = None
            for i, p in enumerate(doc.paragraphs):
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    if pPr.find(qn('w:sectPr')) is not None:
                        break_para_idx = i
                        break

            if break_para_idx is not None:
                first_section_paras = list(range(0, break_para_idx + 1))
            else:
                # Fallback: assume first 4 paragraphs
                first_section_paras = list(range(0, min(4, num_paras)))

            yellow_count = 0
            total_first = len(first_section_paras)
            target_color = 'FFFFF0'
            for idx in first_section_paras:
                fill = get_para_shading_fill(doc.paragraphs[idx])
                if fill and color_distance_hex(fill.upper(), target_color) < 30:
                    yellow_count += 1

            if total_first > 0 and yellow_count == total_first:
                print(f"PASS: Component 5 — All {total_first} first-section paragraphs have yellow background (0.20 pts)")
                total_score += 0.20
            elif total_first > 0 and yellow_count > 0:
                partial = 0.20 * (yellow_count / total_first)
                print(f"PARTIAL: Component 5 — {yellow_count}/{total_first} paras have yellow background ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — No first-section paragraphs have yellow (#FFFFF0) background")
        else:
            print(f"FAIL: Component 5 — Cannot identify first section paragraphs with only {num_sections} section(s)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Second section paragraphs have light blue (#F0F8FF) background (0.20 points)
    # Initial has no paragraph shading. Golden has F0F8FF on paras 4-10.
    try:
        if num_sections >= 2:
            # Second section paragraphs: after the section break
            if break_para_idx is not None:
                second_section_paras = list(range(break_para_idx + 1, num_paras))
            else:
                second_section_paras = list(range(4, num_paras))

            blue_count = 0
            total_second = len(second_section_paras)
            target_color_blue = 'F0F8FF'
            for idx in second_section_paras:
                fill = get_para_shading_fill(doc.paragraphs[idx])
                if fill and color_distance_hex(fill.upper(), target_color_blue) < 30:
                    blue_count += 1

            if total_second > 0 and blue_count == total_second:
                print(f"PASS: Component 6 — All {total_second} second-section paragraphs have blue background (0.20 pts)")
                total_score += 0.20
            elif total_second > 0 and blue_count > 0:
                partial = 0.20 * (blue_count / total_second)
                print(f"PARTIAL: Component 6 — {blue_count}/{total_second} paras have blue background ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — No second-section paragraphs have blue (#F0F8FF) background")
        else:
            print(f"FAIL: Component 6 — Cannot identify second section paragraphs with only {num_sections} section(s)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice edits
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
