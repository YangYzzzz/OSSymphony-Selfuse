"""
Initial Setup: Brochure document with introductory and feature text, no sections defined.
Task ID: writer_rd_068
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_068'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document title ---
    title = doc.add_heading('GreenLeaf Smart Home — Product Brochure', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # === Introductory paragraphs (3 paragraphs) ===
    intro1 = doc.add_paragraph()
    run = intro1.add_run(
        'Welcome to GreenLeaf Smart Home, the next generation of intelligent living. '
        'Our mission is to transform ordinary residences into connected, energy-efficient '
        'sanctuaries that adapt to your lifestyle. Since our founding in 2019, we have '
        'helped over 50,000 households across North America embrace a smarter way of living.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    intro2 = doc.add_paragraph()
    run = intro2.add_run(
        'The GreenLeaf ecosystem integrates seamlessly with major voice assistants, '
        'including Amazon Alexa, Google Home, and Apple HomeKit. Our proprietary '
        'mesh-networking protocol ensures reliable connectivity even in large homes, '
        'covering up to 5,000 square feet without dead zones. Every device is designed '
        'with privacy first — all processing happens locally on your GreenLeaf Hub.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    intro3 = doc.add_paragraph()
    run = intro3.add_run(
        'Whether you are building a new home or retrofitting an existing one, GreenLeaf '
        'offers flexible installation packages tailored to your needs. Our certified '
        'technicians complete most setups in under four hours, and our 24/7 customer '
        'support team is always ready to assist. Read on to discover the features that '
        'make GreenLeaf the preferred choice for modern homeowners.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # === Feature description paragraphs (6 paragraphs) ===
    feat_heading = doc.add_heading('Key Features', level=2)

    feat1 = doc.add_paragraph()
    run = feat1.add_run(
        'Adaptive Climate Control: The GreenLeaf Thermostat learns your daily routine '
        'within the first week of use. It adjusts heating and cooling schedules automatically, '
        'reducing energy consumption by an average of 23%. Zone-based temperature management '
        'lets you set different comfort levels for bedrooms, living areas, and home offices.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    feat2 = doc.add_paragraph()
    run = feat2.add_run(
        'Intelligent Lighting System: Our LED smart bulbs support 16 million colors and '
        'adjustable color temperatures ranging from 2700K warm white to 6500K daylight. '
        'Create custom scenes like "Movie Night," "Morning Energize," or "Dinner Ambiance" '
        'and trigger them with a single voice command or scheduled automation.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    feat3 = doc.add_paragraph()
    run = feat3.add_run(
        'Advanced Security Suite: The GreenLeaf security package includes 4K HDR cameras '
        'with night vision, smart door locks with biometric authentication, and motion-activated '
        'floodlights. All footage is stored locally on an encrypted 2TB drive, with optional '
        'cloud backup for an additional $4.99 per month.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    feat4 = doc.add_paragraph()
    run = feat4.add_run(
        'Water Management System: Monitor your entire household water usage in real time. '
        'The GreenLeaf Water Sensor detects leaks within seconds and automatically shuts off '
        'the main valve to prevent damage. Homeowners report saving an average of 15% on their '
        'annual water bills after installation.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    feat5 = doc.add_paragraph()
    run = feat5.add_run(
        'Energy Dashboard & Solar Integration: The GreenLeaf app provides a comprehensive '
        'energy dashboard showing real-time consumption by device, room, and time of day. '
        'For homes with solar panels, the system optimizes battery storage and grid feed-in '
        'to maximize your return on investment.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    feat6 = doc.add_paragraph()
    run = feat6.add_run(
        'Whole-Home Audio & Entertainment: Stream music to any room through GreenLeaf '
        'wireless speakers with multi-room synchronization. Built-in Chromecast and AirPlay 2 '
        'support means you can cast content from any device. The system also integrates with '
        'your smart TV to create immersive surround sound experiences without additional wiring.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # Save — no sections, no columns, no backgrounds
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
