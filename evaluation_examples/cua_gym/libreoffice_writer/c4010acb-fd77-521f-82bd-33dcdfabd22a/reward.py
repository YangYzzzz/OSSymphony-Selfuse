"""
Reward Script: Conditional text with SET/IF fields for Chapter/Chapters toggling
Task ID: writer_acad_075
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30) - SET field exists defining a variable (e.g. ChapterCount)
  Component 2 (0.40) - IF conditional fields exist that toggle between "Chapter" and "Chapters"
  Component 3 (0.15) - Input instruction paragraph present explaining how to toggle
  Component 4 (0.15) - At least 3 distinct paragraphs contain IF conditional fields
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_075'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verifying."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    body = doc.element.body

    # Collect all instrText elements from the document
    all_instrs = body.findall('.//w:instrText', ns)
    instr_texts = [el.text.strip() if el.text else '' for el in all_instrs]
    print(f"INFO: Found {len(instr_texts)} instrText elements in document")
    for t in instr_texts:
        print(f"  instrText: {repr(t)}")

    # Component 1: SET field exists defining a variable (0.30 points)
    # The golden file should have a SET field like: SET ChapterCount "single"
    # The initial file has NO field codes at all.
    try:
        set_fields = [t for t in instr_texts if t.upper().startswith('SET ')]
        if len(set_fields) > 0:
            # Verify the SET field references Chapter-related toggling
            has_chapter_set = any(
                'chapter' in t.lower() or 'single' in t.lower() or 'multiple' in t.lower()
                for t in set_fields
            )
            if has_chapter_set:
                print(f"PASS: Component 1 - SET field found for chapter variable: {set_fields} (0.30 pts)")
                total_score += 0.30
            else:
                print(f"PARTIAL: Component 1 - SET field found but doesn't reference chapter/single/multiple: {set_fields}")
                total_score += 0.15
        else:
            print(f"FAIL: Component 1 - No SET field found. Expected a SET field for chapter count variable.")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: IF conditional fields exist toggling "Chapter" vs "Chapters" (0.40 points)
    # The golden file should have IF fields like: IF ChapterCount = "single" "Chapter" "Chapters"
    # The initial file has NO IF fields.
    try:
        if_fields = [t for t in instr_texts if t.upper().startswith('IF ')]
        chapter_if_fields = [
            t for t in if_fields
            if ('chapter' in t.lower() and 'chapters' in t.lower())
            or ('"chapter"' in t.lower() and '"chapters"' in t.lower())
        ]
        if len(chapter_if_fields) >= 5:
            print(f"PASS: Component 2 - {len(chapter_if_fields)} IF conditional fields for Chapter/Chapters (0.40 pts)")
            total_score += 0.40
        elif len(chapter_if_fields) >= 2:
            partial = 0.20
            print(f"PARTIAL: Component 2 - {len(chapter_if_fields)} IF fields (need >=5 for full credit) ({partial} pts)")
            total_score += partial
        elif len(chapter_if_fields) >= 1:
            partial = 0.10
            print(f"PARTIAL: Component 2 - Only {len(chapter_if_fields)} IF field found ({partial} pts)")
            total_score += partial
        else:
            # Check for any IF fields at all (maybe different wording)
            if len(if_fields) > 0:
                print(f"FAIL: Component 2 - Found {len(if_fields)} IF fields but none toggle Chapter/Chapters: {if_fields[:3]}")
            else:
                print(f"FAIL: Component 2 - No IF conditional fields found at all.")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Input instruction/reference mode paragraph present (0.15 points)
    # The golden file adds a paragraph explaining how to change the input field value
    # to toggle between single/multiple. Initial file does NOT have this.
    try:
        found_input_instruction = False
        for para in doc.paragraphs:
            text_lower = para.text.lower()
            # Check for instruction text about toggling or input field
            if (('single' in text_lower and 'multiple' in text_lower) or
                ('input field' in text_lower) or
                ('reference mode' in text_lower and ('single' in text_lower or 'chapter' in text_lower))):
                # Also verify this paragraph contains the SET field (not just plain text)
                para_instrs = para._element.findall('.//w:instrText', ns)
                para_flds = para._element.findall('.//w:fldChar', ns)
                if len(para_flds) > 0 or ('single' in text_lower and 'multiple' in text_lower):
                    found_input_instruction = True
                    print(f"PASS: Component 3 - Input instruction paragraph found: '{para.text[:80]}...' (0.15 pts)")
                    break
        if found_input_instruction:
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 - No input instruction paragraph found for toggling chapter references")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: IF fields spread across at least 3 distinct paragraphs (0.15 points)
    # This ensures the conditional text is applied throughout the document, not just once.
    # Initial file has no IF fields in any paragraph.
    try:
        paras_with_if = 0
        for para in doc.paragraphs:
            para_instrs = para._element.findall('.//w:instrText', ns)
            for instr in para_instrs:
                if instr.text and instr.text.strip().upper().startswith('IF '):
                    paras_with_if += 1
                    break  # count each paragraph only once
        if paras_with_if >= 3:
            print(f"PASS: Component 4 - IF fields found in {paras_with_if} distinct paragraphs (>=3) (0.15 pts)")
            total_score += 0.15
        elif paras_with_if >= 1:
            partial = 0.05
            print(f"PARTIAL: Component 4 - IF fields in only {paras_with_if} paragraphs (need >=3) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - No paragraphs contain IF conditional fields")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
