"""
Initial Setup: Thesis document needing conditional text for chapter references
Task ID: writer_acad_075
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_075'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # -- Title Page --
    title = doc.add_heading('', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Dynamic Environmental Modeling\nfor Urban Heat Island Mitigation')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = subtitle.add_run('Master of Science Thesis')
    sr.font.size = Pt(16)
    sr.italic = True

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    ar = author.add_run('Elena Vasquez\nDepartment of Environmental Engineering\nStanford University\nMarch 2026')
    ar.font.size = Pt(12)

    doc.add_page_break()

    # -- Abstract --
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This thesis presents a comprehensive analysis of urban heat island (UHI) '
        'effects in metropolitan areas with populations exceeding 2 million. Through '
        'satellite thermal imaging, ground-based sensor networks, and computational '
        'fluid dynamics simulations, we quantify the temperature differential between '
        'urban cores and surrounding rural areas across 14 cities on 4 continents. '
        'Our findings indicate an average temperature elevation of 3.7°C during summer '
        'months, with peak differentials reaching 8.2°C in densely built environments.'
    )
    doc.add_paragraph(
        'The mitigation strategies evaluated include green roof installations, '
        'reflective surface coatings, urban tree canopy expansion, and permeable '
        'pavement systems. Results demonstrate that combined approaches reduce peak '
        'UHI intensity by 42–58%, depending on local climate and urban morphology.'
    )

    doc.add_page_break()

    # -- Table of Contents placeholder --
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('Chapter 1: Introduction', '1'),
        ('Chapter 2: Literature Review', '8'),
        ('Chapter 3: Methodology', '24'),
        ('Chapter 4: Data Collection and Processing', '41'),
        ('Chapter 5: Results and Analysis', '63'),
        ('Chapter 6: Discussion', '89'),
        ('Chapter 7: Conclusions and Future Work', '104'),
        ('References', '112'),
        ('Appendices', '125'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run(f'\t{page}')

    doc.add_page_break()

    # -- Chapter 1: Introduction --
    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph(
        'Urban heat islands represent one of the most well-documented phenomena in '
        'urban climatology. First identified by Luke Howard in the early 19th century, '
        'the effect has intensified with rapid urbanization worldwide. As of 2025, '
        'approximately 56% of the global population resides in urban areas, a figure '
        'projected to reach 68% by 2050 (United Nations, 2024).'
    )
    doc.add_paragraph(
        'The thermal characteristics of urban environments differ markedly from their '
        'rural counterparts due to several factors: the replacement of vegetation with '
        'impervious surfaces, anthropogenic heat generation, altered wind patterns '
        'caused by building geometry, and reduced evapotranspiration. These factors '
        'combine to create persistent temperature elevations that affect energy '
        'consumption, air quality, public health, and ecological systems.'
    )

    doc.add_heading('1.2 Research Objectives', level=2)
    doc.add_paragraph(
        'This research addresses the following objectives, detailed across multiple '
        'chapters of this thesis:'
    )
    objectives = [
        'Quantify UHI intensity variations across different urban morphologies',
        'Evaluate the effectiveness of green infrastructure interventions',
        'Develop a predictive model for UHI intensity based on urban form parameters',
        'Propose evidence-based mitigation guidelines for urban planners',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Number')

    # Plain text references to chapters (NO conditional text yet - that's the task)
    doc.add_paragraph(
        'The methodology for data collection is described in Chapter 3, while '
        'Chapter 4 covers the processing pipeline. Results are presented in '
        'Chapter 5 and discussed in Chapter 6.'
    )

    doc.add_heading('1.3 Scope and Limitations', level=2)
    doc.add_paragraph(
        'This study focuses on mid-latitude cities in temperate and subtropical '
        'climates. Tropical and arid urban environments, while exhibiting significant '
        'UHI effects, present distinct thermodynamic characteristics that warrant '
        'separate investigation. The analysis covers the period from January 2020 to '
        'December 2025, utilizing Landsat-9 and Sentinel-3 thermal imagery.'
    )

    # A cross-reference paragraph that would benefit from conditional text
    doc.add_paragraph(
        'For a complete overview, see Chapters 3 through 5, which together describe '
        'the full experimental workflow. When referencing a single methodology step, '
        'consult the relevant Chapter individually.'
    )

    doc.add_page_break()

    # -- Chapter 2: Literature Review --
    doc.add_heading('Chapter 2: Literature Review', level=1)

    doc.add_heading('2.1 Historical Context', level=2)
    doc.add_paragraph(
        'The study of urban climatology has evolved significantly since Howard\'s '
        'pioneering observations in London (1818). Oke\'s seminal work (1973, 1982) '
        'established the theoretical framework for understanding urban-rural temperature '
        'contrasts, introducing the distinction between urban canopy layer (UCL) and '
        'urban boundary layer (UBL) heat islands.'
    )

    doc.add_heading('2.2 Remote Sensing Approaches', level=2)
    doc.add_paragraph(
        'Satellite-based thermal remote sensing has revolutionized UHI research by '
        'enabling city-scale temperature mapping. Voogt and Oke (2003) provided a '
        'comprehensive review of thermal remote sensing in urban areas, highlighting '
        'the distinction between surface and air temperature measurements. More recent '
        'work by Zhou et al. (2019) demonstrated the utility of combining Landsat and '
        'MODIS data for multi-temporal UHI analysis.'
    )

    doc.add_heading('2.3 Mitigation Strategies', level=2)
    doc.add_paragraph(
        'Green roofs have emerged as a leading UHI mitigation strategy. Research by '
        'Santamouris (2014) found that extensive green roof installations can reduce '
        'ambient air temperature by 0.3–3.0°C at the neighborhood scale. Akbari et al. '
        '(2001) demonstrated that increasing urban albedo through reflective surfaces '
        'can offset 1–2°C of UHI-induced warming.'
    )

    # Cross-reference that mentions chapters
    doc.add_paragraph(
        'The mitigation strategies reviewed here are evaluated quantitatively in '
        'Chapter 5. Additional technical details on measurement protocols can be found '
        'in Chapter 3.'
    )

    doc.add_page_break()

    # -- Chapter 3: Methodology (abbreviated) --
    doc.add_heading('Chapter 3: Methodology', level=1)

    doc.add_heading('3.1 Study Area Selection', level=2)
    doc.add_paragraph(
        'Fourteen cities were selected based on the following criteria: population '
        'exceeding 2 million, availability of high-resolution thermal satellite imagery, '
        'presence of ground-based weather stations within the urban core, and '
        'representation of diverse climate zones (Köppen classification Cfa, Cfb, Csa, '
        'and Cwa).'
    )

    # Table of study cities
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    headers = ['City', 'Country', 'Population (M)', 'Köppen Zone']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    cities_data = [
        ['New York', 'United States', '8.3', 'Cfa'],
        ['London', 'United Kingdom', '9.0', 'Cfb'],
        ['Tokyo', 'Japan', '13.9', 'Cfa'],
        ['Sydney', 'Australia', '5.3', 'Cfa'],
        ['São Paulo', 'Brazil', '12.3', 'Cwa'],
        ['Barcelona', 'Spain', '5.6', 'Csa'],
        ['Shanghai', 'China', '24.9', 'Cfa'],
    ]
    for r, row_data in enumerate(cities_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacer

    doc.add_heading('3.2 Data Acquisition', level=2)
    doc.add_paragraph(
        'Thermal data was acquired from two primary satellite platforms: Landsat-9 '
        '(100m spatial resolution, 16-day revisit) and Sentinel-3 SLSTR (1km resolution, '
        'daily revisit). Ground truth measurements were obtained from 247 weather '
        'stations distributed across the 14 study cities, providing hourly temperature '
        'readings at 2m height.'
    )

    doc.add_paragraph(
        'Note: The experimental setup described in this chapter is complemented by '
        'the data processing methods in Chapter 4. For the combined analysis across '
        'Chapters 3, 4, and 5, see the integrated results discussion in Chapter 6.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
