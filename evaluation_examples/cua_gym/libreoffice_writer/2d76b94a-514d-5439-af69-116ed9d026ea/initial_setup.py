"""
Initial Setup: Employee directory document with 10 employees
Task ID: writer_hr_015
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Greenfield Solutions — Employee Directory', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Human Resources Department — Confidential')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    doc.add_paragraph('')  # blank line

    # Employee data: (name, department, email)
    employees = [
        ('Sarah Chen', 'Engineering', 's.chen@greenfieldsolutions.com'),
        ('Marcus Johnson', 'Marketing', 'm.johnson@greenfieldsolutions.com'),
        ('Priya Patel', 'Finance', 'p.patel@greenfieldsolutions.com'),
        ('David Kim', 'Human Resources', 'd.kim@greenfieldsolutions.com'),
        ('Elena Rodriguez', 'Operations', 'e.rodriguez@greenfieldsolutions.com'),
        ('James O\'Brien', 'Sales', 'j.obrien@greenfieldsolutions.com'),
        ('Aisha Williams', 'Engineering', 'a.williams@greenfieldsolutions.com'),
        ('Robert Nakamura', 'Legal', 'r.nakamura@greenfieldsolutions.com'),
        ('Lisa Bergström', 'Product Management', 'l.bergstrom@greenfieldsolutions.com'),
        ('Carlos Mendoza', 'Customer Support', 'c.mendoza@greenfieldsolutions.com'),
    ]

    for i, (name, dept, email) in enumerate(employees, 1):
        # Employee name as bold
        para_name = doc.add_paragraph()
        run_num = para_name.add_run(f'{i}. ')
        run_num.font.size = Pt(11)
        run_name = para_name.add_run(name)
        run_name.bold = True
        run_name.font.size = Pt(12)

        # Department
        para_dept = doc.add_paragraph()
        para_dept.paragraph_format.left_indent = Pt(24)
        run_dept_label = para_dept.add_run('Department: ')
        run_dept_label.font.size = Pt(11)
        run_dept_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run_dept_val = para_dept.add_run(dept)
        run_dept_val.font.size = Pt(11)

        # Email - regular formatting, NO underline
        para_email = doc.add_paragraph()
        para_email.paragraph_format.left_indent = Pt(24)
        run_email_label = para_email.add_run('Email: ')
        run_email_label.font.size = Pt(11)
        run_email_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run_email_val = para_email.add_run(email)
        run_email_val.font.size = Pt(11)
        run_email_val.underline = False  # explicitly not underlined

        # Blank line separator (except after last)
        if i < len(employees):
            doc.add_paragraph('')

    # Footer note
    doc.add_paragraph('')
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_footer = footer_para.add_run('Last updated: March 2026')
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_footer.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
