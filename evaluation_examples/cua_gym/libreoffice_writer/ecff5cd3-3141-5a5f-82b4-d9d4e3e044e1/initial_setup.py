"""
Initial Setup: Create a novel master document with 10 chapters, all using default page style.
Task ID: writer_rm_076
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_076'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# Chapter content - realistic novel excerpts
CHAPTERS = [
    {
        "title": "The Awakening",
        "paragraphs": [
            "The morning light crept through the curtains of the small apartment on Maple Street, casting long shadows across the hardwood floor. Elena Vasquez sat at her kitchen table, nursing a cup of coffee that had long gone cold, staring at the letter that had arrived the previous evening.",
            "It bore no return address, only her name written in an elegant script she didn't recognize. The paper was thick, cream-colored, the kind used for formal invitations or important correspondence. She had read it three times already, and each time the words seemed more impossible than the last.",
            "\"Dear Ms. Vasquez,\" it began. \"We regret to inform you that your grandmother, Isabela Montero de Vasquez, has passed away at her estate in Andalusia. As her sole living heir, you are requested to attend the reading of her will at the offices of Delgado & Partners, Seville, no later than the 15th of next month.\"",
            "Elena had never known her grandmother. Her father, Rodrigo, had left Spain before Elena was born and rarely spoke of his family. When pressed, he would say only that some doors were better left closed. Now, with both her parents gone, this letter was a bridge to a past she had never been allowed to explore.",
        ],
    },
    {
        "title": "The Journey South",
        "paragraphs": [
            "The flight from New York to Madrid took seven hours, during which Elena managed to sleep only fitfully. She had packed light — a single carry-on and her laptop bag — unsure of how long she would need to stay. The law firm's follow-up email had been vague about the timeline.",
            "At Barajas Airport, she rented a small Seat Ibiza and programmed the GPS for Seville, a drive of roughly five hours through the heart of Spain. The landscape shifted as she drove south: the sprawling suburbs of Madrid gave way to the vast, sun-scorched plains of La Mancha.",
            "She stopped for lunch at a roadside venta near Valdepeñas, ordering a plate of migas and a glass of local red wine. The owner, a stout woman named Pilar, asked where she was headed. When Elena mentioned the Vasquez estate outside Seville, the woman's expression changed almost imperceptibly.",
            "\"Ah, the Finca de los Naranjos,\" Pilar said, wiping her hands on her apron. \"Everyone in the south knows that place. Beautiful, but... be careful, señorita. Old houses have long memories.\"",
        ],
    },
    {
        "title": "Finca de los Naranjos",
        "paragraphs": [
            "The estate appeared at the end of a long, unpaved road lined with orange trees, their branches heavy with fruit that glowed in the late afternoon sun. The main house was a sprawling cortijo of whitewashed walls and terracotta tiles, its facade partly covered in bougainvillea.",
            "A man in his sixties waited at the entrance. He introduced himself as Tomás Delgado, senior partner at the law firm. He was tall and thin, with silver hair combed neatly back, and wore a linen suit despite the October heat.",
            "\"Thank you for coming, Señorita Vasquez,\" he said, shaking her hand. \"Your grandmother specified in her will that the reading must take place here, at the finca. She was very particular about that.\"",
            "Elena followed him through an arched doorway into a courtyard filled with potted jasmine and a central fountain that no longer ran. The house smelled of old wood, dried lavender, and something else — something faintly metallic that she couldn't quite place.",
            "Delgado led her to a study on the ground floor. The walls were lined with bookshelves that reached the ceiling, filled with leather-bound volumes in Spanish, French, and Arabic. A large mahogany desk dominated the center of the room.",
        ],
    },
    {
        "title": "The Will",
        "paragraphs": [
            "Delgado opened a leather portfolio and removed a document sealed with red wax. He broke the seal carefully and unfolded the pages. The will was handwritten, in the same elegant script as the letter Elena had received in New York.",
            "\"I, Isabela Montero de Vasquez, being of sound mind, do hereby bequeath the entirety of my estate — the Finca de los Naranjos, its grounds, and all contents therein — to my granddaughter, Elena Sofia Vasquez, on the condition that she resides at the property for no fewer than ninety consecutive days following the reading of this will.\"",
            "Elena stared at the lawyer. \"Ninety days? I have a job. I have a life in New York.\"",
            "Delgado nodded sympathetically. \"Your grandmother anticipated that reaction. She also left this.\" He handed Elena a smaller envelope. Inside was a handwritten note: \"Elena, I know you have questions. The answers are in the house. Give it ninety days. — Your Abuela, Isabela.\"",
            "There was also a bank statement showing an account in Elena's name with a balance of four hundred and seventy-eight thousand euros. \"To cover any expenses during the residency period,\" Delgado explained.",
        ],
    },
    {
        "title": "Settling In",
        "paragraphs": [
            "The housekeeper, a quiet woman named Carmen who had served Isabela for thirty years, showed Elena to the master bedroom on the second floor. The room was large, with a four-poster bed draped in white linen and tall windows overlooking the orange groves.",
            "Carmen had maintained the house impeccably. Every surface was clean, the linens fresh, the kitchen stocked. It was as if Isabela had simply stepped out and might return at any moment.",
            "Over the next few days, Elena began to explore the property. The finca comprised the main house, a smaller guest cottage, a chapel, stables that now housed only dust and swallows, and nearly twenty hectares of orange and olive groves.",
            "She called her employer, a digital marketing firm called BrightPath Media, and requested a leave of absence. Her manager, David Chen, was reluctant but agreed to three months, noting that Elena's projects could be reassigned temporarily.",
            "She also called her best friend, Nadia Okonkwo, who was predictably outraged. \"You inherited an estate in Spain and you're thinking of NOT going? Elena, this is the plot of every good novel. Go live it.\"",
        ],
    },
    {
        "title": "The Library",
        "paragraphs": [
            "On her fifth day at the finca, Elena turned her attention to the library. It was the room that drew her most — perhaps because of the sheer volume of knowledge it contained, or perhaps because it was clearly the room Isabela had loved most.",
            "The collection was extraordinary. First editions of Cervantes and Lorca sat alongside treatises on botany, astronomy, and what appeared to be alchemy. There were journals in Isabela's handwriting dating back to the 1960s, filled with observations about the natural world, sketches of plants, and recipes for herbal remedies.",
            "Behind a section of botanical texts, Elena found a locked drawer in the wall. Carmen, when asked, said she knew nothing about it. \"Your grandmother had her secrets,\" she said simply.",
            "Elena searched the desk and found a small brass key hidden inside a hollowed-out copy of Don Quixote. It fit the drawer perfectly. Inside was a leather journal, older than the others, its pages yellowed with age.",
            "The first entry was dated March 12, 1952, and written not by Isabela, but by someone named Rafael Montero — Elena's great-grandfather.",
        ],
    },
    {
        "title": "Rafael's Journal",
        "paragraphs": [
            "Rafael Montero had been a professor of medieval history at the University of Seville before the Civil War. His journal described years of research into the Moorish occupation of southern Spain, specifically the period between 1100 and 1248.",
            "He wrote extensively about a rumored repository of knowledge — a hidden library established by the scholars of the Almohad Caliphate, said to contain texts that had been translated from Greek, Persian, and Sanskrit. The location, according to Rafael, was somewhere in the hills east of Seville.",
            "\"The local legends speak of a cave system near the village of Constantina,\" Rafael wrote. \"The shepherds avoid it, claiming the hills are haunted. But I have found references in the university archives that suggest the caves were used by Almohad scholars as a retreat during the Reconquista.\"",
            "The journal entries grew increasingly excited over the following months. Rafael described finding pottery shards, fragments of Arabic manuscripts, and architectural remnants that supported his theory. Then, abruptly, the entries stopped.",
            "The last entry, dated November 3, 1952, read simply: \"I have found it. God help me, I have found it. I must speak with Isabela before I proceed.\"",
        ],
    },
    {
        "title": "The Cave",
        "paragraphs": [
            "Armed with Rafael's journal and a topographic map of the region, Elena drove to Constantina, a small white village nestled in the Sierra Norte. She parked near the church square and walked to a bar called El Mirador, where she ordered a café con leche and asked the bartender about caves in the area.",
            "The bartender, a young man named Javier, confirmed that there were caves in the hills. \"The old folks say they go deep into the mountain. Some explorers went in during the eighties and found Moorish tiles. But the entrance collapsed in a storm a few years later.\"",
            "Elena spent two days hiking the hills east of the village, following the landmarks described in Rafael's journal — a split rock shaped like a bull's head, a dry riverbed that curved sharply westward, an ancient cork oak with a trunk wide enough to shelter three people.",
            "On the second evening, as the setting sun cast the landscape in shades of gold and copper, she found it: a narrow opening in the hillside, partially concealed by scrub brush and fallen stone. The collapse Javier had mentioned was only partial — the entrance was tight, but passable.",
            "She squeezed through with her flashlight and descended into cool, still air that smelled of mineral water and ancient dust.",
        ],
    },
    {
        "title": "The Discovery",
        "paragraphs": [
            "The cave opened into a series of chambers, their walls smooth and marked with geometric patterns carved into the stone. The craftsmanship was unmistakable — this was not a natural formation. Someone had shaped these spaces with purpose and skill.",
            "In the third chamber, Elena found what Rafael had described. Stone shelves lined the walls, and on them sat dozens of clay cylinders, each containing tightly rolled manuscripts. The preservation was remarkable — the cave's constant temperature and low humidity had protected the documents for nearly a millennium.",
            "She photographed everything methodically, using her phone's camera and a small LED panel she had brought for light. She did not touch the manuscripts — she had enough sense to know that untrained handling could destroy them.",
            "Back at the finca, she called Professor Amira Hassan at the University of Granada, a specialist in Andalusian Islamic history whom she had found through an online search. When Elena described her find, Hassan was silent for a long moment.",
            "\"If what you're describing is real,\" Hassan said carefully, \"this could be one of the most significant archaeological finds in Spanish history. I'll assemble a team and come to you. Don't tell anyone else. Not yet.\"",
        ],
    },
    {
        "title": "The Decision",
        "paragraphs": [
            "Professor Hassan arrived three days later with a team of four: two archaeologists, a conservator, and a linguist specializing in medieval Arabic. They spent a week documenting the cave and carefully extracting sample manuscripts for preliminary analysis.",
            "The initial findings were staggering. The collection appeared to include translations of lost works by Aristotle, original medical treatises from the Baghdad House of Wisdom, and astronomical tables that predated Copernicus by three centuries.",
            "Hassan sat with Elena on the terrace of the finca as the sun set behind the orange groves. \"Your great-grandfather was right,\" she said. \"And your grandmother must have known. That's why she wanted you here — to finish what Rafael started.\"",
            "Elena thought of her grandmother's note: \"The answers are in the house.\" Isabela had kept the secret for seventy years, waiting for someone she trusted to carry it forward. She had chosen Elena.",
            "Elena looked out at the groves, the hills beyond, the fading light. In New York, her old life waited — the apartment, the job, the routines. But here, in this house full of memories and mysteries, she had found something she hadn't known she was looking for: a purpose, a history, a home.",
            "She picked up her phone and called David Chen. \"I won't be coming back,\" she said. \"I've found where I need to be.\"",
        ],
    },
]


def create_initial():
    doc = Document()

    # Set default page style - standard margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title page
    title_para = doc.add_heading("The Orange Grove Legacy", level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run("by Elena Sofia Vasquez")
    run.font.size = Pt(14)
    run.font.italic = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle_para.add_run("A Novel")
    run.font.size = Pt(12)

    # Add page break after title page
    doc.add_page_break()

    # Add each chapter as a new section
    for i, chapter in enumerate(CHAPTERS):
        if i > 0:
            # Add section break for chapters 2-10
            new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            new_section.top_margin = Inches(1)
            new_section.bottom_margin = Inches(1)
            new_section.left_margin = Inches(1.25)
            new_section.right_margin = Inches(1.25)

        # Chapter heading
        heading = doc.add_heading(f"Chapter {i + 1}: {chapter['title']}", level=1)

        # Chapter paragraphs
        for text in chapter["paragraphs"]:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Inches(0.5)

    # Add a simple uniform header and footer to ALL sections (no first-page differentiation)
    for section in doc.sections:
        # Header - same on all pages
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = "The Orange Grove Legacy"
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.italic = True

        # Footer - same on all pages (page number)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Page number field
        r1 = fp.add_run()
        r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
        r2 = fp.add_run()
        instr = r2._element.makeelement(qn('w:instrText'), {})
        instr.text = ' PAGE '
        r2._element.append(instr)
        r3 = fp.add_run()
        r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

        # IMPORTANT: No different first page - this is what the task asks the agent to change
        section.different_first_page_header_footer = False

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
