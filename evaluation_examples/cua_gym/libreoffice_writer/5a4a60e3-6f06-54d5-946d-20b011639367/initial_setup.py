"""
Initial Setup: Annual Budget Report - no header/footer (pre-task state)
Task ID: writer_page_065
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
DESKTOP = '/home/user/Desktop'
TASK_ID = 'writer_page_065'
OUTPUT = f'{DESKTOP}/budget_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1, bold=True, size=14):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if level == 0:
        run.font.size = Pt(24)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif level == 1:
        run.font.size = Pt(16)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
    return para


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    return para


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set page size to A4, portrait, 2.54cm margins on all sides
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Ensure no header/footer on the section
    section.different_first_page_header_footer = False
    # Leave header and footer empty (do not add content)

    # -----------------------------------------------------------------------
    # PAGE 1: Cover Page
    # -----------------------------------------------------------------------
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(72)
    run = para.add_run("ANNUAL BUDGET REPORT")
    run.bold = True
    run.font.size = Pt(28)

    para2 = doc.add_paragraph()
    para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para2.paragraph_format.space_before = Pt(12)
    run2 = para2.add_run("Fiscal Year 2025")
    run2.bold = True
    run2.font.size = Pt(18)

    para3 = doc.add_paragraph()
    para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para3.paragraph_format.space_before = Pt(6)
    run3 = para3.add_run("Meridian Technologies Group")
    run3.font.size = Pt(14)

    para4 = doc.add_paragraph()
    para4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para4.paragraph_format.space_before = Pt(6)
    run4 = para4.add_run("Prepared by: Finance Department")
    run4.font.size = Pt(12)

    para5 = doc.add_paragraph()
    para5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para5.paragraph_format.space_before = Pt(4)
    run5 = para5.add_run("Date: January 15, 2025")
    run5.font.size = Pt(12)

    para6 = doc.add_paragraph()
    para6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para6.paragraph_format.space_before = Pt(4)
    run6 = para6.add_run("Confidential — For Internal Use Only")
    run6.italic = True
    run6.font.size = Pt(10)

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 2: Table of Contents
    # -----------------------------------------------------------------------
    add_heading(doc, "Table of Contents", level=1)

    toc_items = [
        ("1. Executive Summary", "3"),
        ("2. Revenue Overview", "4"),
        ("3. Operating Expenses", "5"),
        ("4. Capital Expenditures", "6"),
        ("5. Departmental Budgets", "7"),
        ("6. Variance Analysis", "8"),
        ("7. Forecast & Projections", "9"),
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        run_t = p.add_run(title)
        run_t.font.size = Pt(11)
        # Add dots and page number (simplified)
        run_p = p.add_run(f"{'.' * (60 - len(title))}{page}")
        run_p.font.size = Pt(11)

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 3: Executive Summary
    # -----------------------------------------------------------------------
    add_heading(doc, "1. Executive Summary", level=1)

    doc.add_paragraph(
        "This Annual Budget Report for Fiscal Year 2025 presents a comprehensive financial "
        "analysis of Meridian Technologies Group. The company achieved strong performance "
        "across all major business units, with total revenue reaching $142.7 million, "
        "representing a 12.4% increase over FY2024."
    )
    doc.add_paragraph(
        "Operating expenses totaled $98.3 million, yielding an operating margin of 31.1%. "
        "Net income after taxes stood at $29.8 million, up from $24.1 million in the "
        "prior year. Capital expenditures were maintained at $18.5 million, primarily "
        "directed toward infrastructure modernization and product development initiatives."
    )

    add_heading(doc, "Key Highlights", level=2)
    highlights = [
        "Total Revenue: $142.7M (+12.4% YoY)",
        "Operating Income: $44.4M (31.1% margin)",
        "Net Income: $29.8M (+23.7% YoY)",
        "EBITDA: $52.1M (36.5% margin)",
        "Capital Expenditures: $18.5M",
        "Headcount: 1,247 full-time employees",
    ]
    for h in highlights:
        doc.add_paragraph(h, style="List Bullet")

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 4: Revenue Overview
    # -----------------------------------------------------------------------
    add_heading(doc, "2. Revenue Overview", level=1)

    doc.add_paragraph(
        "Total revenue for FY2025 was $142.7 million, driven by growth in cloud services "
        "and professional consulting segments. Software licensing revenue grew by 8.2%, "
        "while managed services expanded by 21.5% as enterprise clients continued to shift "
        "workloads to cloud-based platforms."
    )

    add_heading(doc, "Revenue by Business Segment", level=2)

    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers_row = table.rows[0]
    for i, h in enumerate(["Segment", "FY2024 ($M)", "FY2025 ($M)", "Growth (%)"]):
        cell = headers_row.cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    data_rows = [
        ["Cloud Services", "38.2", "46.4", "+21.5%"],
        ["Software Licensing", "42.1", "45.5", "+8.2%"],
        ["Professional Services", "26.8", "31.9", "+19.0%"],
        ["Hardware Sales", "14.4", "12.7", "-11.8%"],
        ["Support & Maintenance", "5.3", "6.2", "+17.0%"],
    ]
    for ri, row_data in enumerate(data_rows, 1):
        for ci, val in enumerate(row_data):
            table.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph("")

    doc.add_paragraph(
        "North America remains the primary revenue driver at 68% of total revenue ($97.0M), "
        "followed by the EMEA region at 22% ($31.4M) and Asia-Pacific at 10% ($14.3M). "
        "The company's international expansion strategy is expected to increase APAC revenue "
        "to 15% of total by FY2026."
    )

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 5: Operating Expenses
    # -----------------------------------------------------------------------
    add_heading(doc, "3. Operating Expenses", level=1)

    doc.add_paragraph(
        "Total operating expenses for FY2025 were $98.3 million, an increase of 9.1% from "
        "$90.1 million in FY2024. Cost of revenue was $52.4 million (36.7% of revenue), "
        "while operating expenses excluding COGS were $45.9 million."
    )

    add_heading(doc, "Expense Breakdown", level=2)

    table2 = doc.add_table(rows=7, cols=3)
    table2.style = "Table Grid"
    for i, h in enumerate(["Category", "FY2025 ($M)", "% of Revenue"]):
        run = table2.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    expense_data = [
        ["Cost of Revenue", "52.4", "36.7%"],
        ["Research & Development", "18.7", "13.1%"],
        ["Sales & Marketing", "14.3", "10.0%"],
        ["General & Administrative", "7.8", "5.5%"],
        ["Depreciation & Amortization", "3.6", "2.5%"],
        ["Other Operating Expenses", "1.5", "1.1%"],
    ]
    for ri, row_data in enumerate(expense_data, 1):
        for ci, val in enumerate(row_data):
            table2.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph(
        "R&D investment of $18.7M reflects our commitment to innovation, representing a "
        "significant allocation toward next-generation AI-powered analytics products. "
        "Headcount-related costs account for 68% of total operating expenses."
    )

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 6: Capital Expenditures
    # -----------------------------------------------------------------------
    add_heading(doc, "4. Capital Expenditures", level=1)

    doc.add_paragraph(
        "Capital expenditures in FY2025 totaled $18.5 million, consistent with the approved "
        "budget of $19.0 million. Investments were concentrated in three core areas: "
        "data center infrastructure expansion, enterprise software platforms, and "
        "laboratory equipment for the R&D division."
    )

    add_heading(doc, "CapEx Summary", level=2)

    capex_data = [
        ["Data Center Infrastructure", "7.2", "38.9%"],
        ["Enterprise Software Platforms", "4.8", "25.9%"],
        ["R&D Laboratory Equipment", "3.1", "16.8%"],
        ["Office Facilities", "2.0", "10.8%"],
        ["Vehicles & Fleet", "1.4", "7.6%"],
    ]
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = "Table Grid"
    for i, h in enumerate(["Category", "Amount ($M)", "% of Total CapEx"]):
        run = table3.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for ri, row_data in enumerate(capex_data, 1):
        for ci, val in enumerate(row_data):
            table3.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph(
        "The data center expansion project, which accounts for the largest CapEx allocation, "
        "was completed in Q3 FY2025 and is expected to reduce infrastructure costs by 14% "
        "annually through improved energy efficiency and server consolidation."
    )

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 7: Departmental Budgets
    # -----------------------------------------------------------------------
    add_heading(doc, "5. Departmental Budgets", level=1)

    doc.add_paragraph(
        "Each department received budget allocations aligned with strategic priorities for "
        "FY2025. The Engineering division received the largest allocation to support product "
        "development and cloud infrastructure initiatives."
    )

    departments = [
        ("Engineering", "32.4", "31.6", "-0.8", "-2.5%"),
        ("Sales & Marketing", "18.6", "19.2", "+0.6", "+3.2%"),
        ("Operations", "14.1", "13.8", "-0.3", "-2.1%"),
        ("Finance & Accounting", "5.2", "5.6", "+0.4", "+7.7%"),
        ("Human Resources", "3.8", "4.1", "+0.3", "+7.9%"),
        ("Legal & Compliance", "2.9", "3.0", "+0.1", "+3.4%"),
        ("Executive Office", "2.3", "2.5", "+0.2", "+8.7%"),
    ]
    table4 = doc.add_table(rows=8, cols=5)
    table4.style = "Table Grid"
    for i, h in enumerate(["Department", "Budget ($M)", "Actual ($M)", "Variance ($M)", "Variance (%)"]):
        run = table4.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for ri, row_data in enumerate(departments, 1):
        for ci, val in enumerate(row_data):
            table4.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 8: Variance Analysis
    # -----------------------------------------------------------------------
    add_heading(doc, "6. Variance Analysis", level=1)

    doc.add_paragraph(
        "Variance analysis for FY2025 reveals overall favorable budget performance. Total "
        "actual spending came in at $98.3 million versus a budgeted $99.7 million, representing "
        "a favorable variance of $1.4 million (1.4%). Revenue exceeded budget by $4.2 million "
        "due to stronger-than-expected performance in cloud services."
    )

    add_heading(doc, "Key Variances by Category", level=2)

    variances = [
        ("Revenue", "+4.2", "Favorable", "Cloud services outperformance"),
        ("COGS", "-1.1", "Unfavorable", "Higher-than-expected support costs"),
        ("R&D Expenses", "+0.8", "Favorable", "Delayed hiring in H1"),
        ("Sales & Marketing", "-0.3", "Unfavorable", "Additional campaign spending"),
        ("G&A", "+0.6", "Favorable", "Reduced travel costs"),
        ("CapEx", "+0.5", "Favorable", "Phased infrastructure rollout"),
    ]
    table5 = doc.add_table(rows=7, cols=4)
    table5.style = "Table Grid"
    for i, h in enumerate(["Category", "Variance ($M)", "Type", "Primary Driver"]):
        run = table5.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for ri, row_data in enumerate(variances, 1):
        for ci, val in enumerate(row_data):
            table5.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)

    doc.add_paragraph("")
    doc.add_paragraph(
        "The unfavorable COGS variance was primarily driven by unplanned support escalations "
        "in Q2, which required additional contractor resources. Corrective actions have been "
        "implemented, and the support cost model has been revised for FY2026 planning."
    )

    add_page_break(doc)

    # -----------------------------------------------------------------------
    # PAGE 9: Forecast & Projections
    # -----------------------------------------------------------------------
    add_heading(doc, "7. Forecast & Projections", level=1)

    doc.add_paragraph(
        "Based on current pipeline strength and market conditions, FY2026 revenue is projected "
        "at $163.4 million, representing 14.5% growth over FY2025. Operating margins are "
        "expected to expand to 33.5% as the company benefits from operating leverage in its "
        "cloud services business."
    )

    add_heading(doc, "Three-Year Financial Projections", level=2)

    projections = [
        ("Total Revenue ($M)", "142.7", "163.4", "189.2"),
        ("Revenue Growth (%)", "12.4%", "14.5%", "15.8%"),
        ("Gross Margin (%)", "63.3%", "65.0%", "66.5%"),
        ("Operating Income ($M)", "44.4", "54.7", "68.2"),
        ("Operating Margin (%)", "31.1%", "33.5%", "36.1%"),
        ("Net Income ($M)", "29.8", "37.6", "47.8"),
        ("EPS ($)", "3.42", "4.18", "5.24"),
        ("CapEx ($M)", "18.5", "22.0", "25.5"),
    ]
    table6 = doc.add_table(rows=9, cols=4)
    table6.style = "Table Grid"
    for i, h in enumerate(["Metric", "FY2025 (Actual)", "FY2026 (Est.)", "FY2027 (Est.)"]):
        run = table6.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for ri, row_data in enumerate(projections, 1):
        for ci, val in enumerate(row_data):
            table6.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)

    doc.add_paragraph("")
    doc.add_paragraph(
        "These projections are based on conservative assumptions and are subject to "
        "macroeconomic conditions, competitive dynamics, and the successful execution "
        "of our strategic initiatives. The Finance team will provide quarterly updates "
        "to the Board of Directors."
    )

    doc.add_paragraph("")
    p_end = doc.add_paragraph()
    p_end.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_end = p_end.add_run("— End of Annual Budget Report FY2025 —")
    run_end.italic = True
    run_end.font.size = Pt(10)

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
