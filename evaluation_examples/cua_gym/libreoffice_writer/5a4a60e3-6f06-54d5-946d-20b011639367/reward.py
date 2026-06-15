"""
Reward Script: Header/Footer setup for budget_report.docx
Task ID: writer_page_065
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.25): different_first_page_header_footer enabled (titlePg set)
  - Component 2 (0.30): Default header (pages 2+) contains 'Annual Budget Report FY2025', left-aligned
  - Component 3 (0.25): Default footer (pages 2+) contains a PAGE field code, right-aligned
  - Component 4 (0.10): First page header is empty
  - Component 5 (0.10): First page footer is empty
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import lxml.etree as etree

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_065'
FILE_PATH = f'{WORKDIR}/budget_report.docx'


def has_page_field(paragraph):
    """Return True if the paragraph contains a PAGE field code (instrText with PAGE)."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for instr in paragraph._element.findall(f'.//{{{ns}}}instrText'):
        if 'PAGE' in (instr.text or ''):
            return True
    return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Gate: file must be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) == 0:
        print("CRITICAL: Document has no sections")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]

    # Component 1: different_first_page_header_footer is enabled (0.25 points)
    # This is indicated by <w:titlePg/> in the sectPr XML
    # In python-docx, section.different_first_page_header_footer checks for titlePg
    try:
        diff_first = section.different_first_page_header_footer
        if diff_first:
            print(f"PASS: Component 1 — different_first_page_header_footer is True (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — different_first_page_header_footer is False (expected True)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Default header (pages 2+) contains 'Annual Budget Report FY2025' left-aligned (0.30 points)
    try:
        hdr = section.header
        # header should not be linked to previous (must have actual content)
        linked = hdr.is_linked_to_previous
        header_text = ""
        header_alignment = None
        if hdr.paragraphs:
            para = hdr.paragraphs[0]
            header_text = para.text.strip()
            header_alignment = para.alignment

        expected_text = "Annual Budget Report FY2025"
        # Alignment LEFT is 0 or WD_PARAGRAPH_ALIGNMENT.LEFT or None (default is left)
        alignment_ok = (
            header_alignment == WD_PARAGRAPH_ALIGNMENT.LEFT or
            header_alignment is None  # default paragraph alignment is left
        )
        text_ok = expected_text in header_text

        if text_ok and alignment_ok and not linked:
            print(f"PASS: Component 2 — Default header contains '{expected_text}' left-aligned (0.30 pts)")
            total_score += 0.30
        elif text_ok and alignment_ok and linked:
            print(f"FAIL: Component 2 — Header text is correct but header is linked to previous")
        elif not text_ok:
            print(f"FAIL: Component 2 — Expected header '{expected_text}', found '{header_text}'")
        else:
            print(f"FAIL: Component 2 — Header alignment wrong: {header_alignment} (expected LEFT or None)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Default footer (pages 2+) has PAGE field code, right-aligned (0.25 points)
    try:
        ftr = section.footer
        footer_linked = ftr.is_linked_to_previous
        footer_has_page_field = False
        footer_alignment = None

        if ftr.paragraphs:
            para = ftr.paragraphs[0]
            footer_has_page_field = has_page_field(para)
            footer_alignment = para.alignment

        # Right alignment is WD_PARAGRAPH_ALIGNMENT.RIGHT = 2
        alignment_right = (footer_alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if footer_has_page_field and alignment_right and not footer_linked:
            print(f"PASS: Component 3 — Default footer has PAGE field code, right-aligned (0.25 pts)")
            total_score += 0.25
        elif footer_has_page_field and not alignment_right:
            print(f"FAIL: Component 3 — Footer has PAGE field but alignment is {footer_alignment} (expected RIGHT)")
        elif not footer_has_page_field:
            print(f"FAIL: Component 3 — Footer missing PAGE field code (footer text: '{ftr.paragraphs[0].text if ftr.paragraphs else ''}'); footer linked={footer_linked}")
        else:
            print(f"FAIL: Component 3 — Footer is still linked to previous (linked={footer_linked})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: First page header is empty AND different_first_page is enabled (0.10 points)
    # This sub-condition only awards credit when the task change (titlePg) is also in place.
    # If different_first_page_header_footer is False (initial state), this cannot pass.
    try:
        diff_first_for_comp4 = section.different_first_page_header_footer
        first_hdr = section.first_page_header
        first_hdr_text = ""
        for para in first_hdr.paragraphs:
            first_hdr_text += para.text.strip()

        if diff_first_for_comp4 and first_hdr_text == "":
            print(f"PASS: Component 4 — different_first_page enabled AND first page header is empty (0.10 pts)")
            total_score += 0.10
        elif not diff_first_for_comp4:
            print(f"FAIL: Component 4 — different_first_page_header_footer not enabled (required for first-page empty header to be meaningful)")
        else:
            print(f"FAIL: Component 4 — First page header is not empty: '{first_hdr_text}'")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: First page footer is empty AND different_first_page is enabled (0.10 points)
    # Same gate: only meaningful when titlePg is set.
    try:
        diff_first_for_comp5 = section.different_first_page_header_footer
        first_ftr = section.first_page_footer
        first_ftr_text = ""
        first_ftr_has_page_field = False

        for para in first_ftr.paragraphs:
            first_ftr_text += para.text.strip()
            first_ftr_has_page_field = first_ftr_has_page_field or has_page_field(para)

        if diff_first_for_comp5 and first_ftr_text == "" and not first_ftr_has_page_field:
            print(f"PASS: Component 5 — different_first_page enabled AND first page footer is empty (0.10 pts)")
            total_score += 0.10
        elif not diff_first_for_comp5:
            print(f"FAIL: Component 5 — different_first_page_header_footer not enabled (required for first-page empty footer to be meaningful)")
        else:
            print(f"FAIL: Component 5 — First page footer is not empty: text='{first_ftr_text}', has_page_field={first_ftr_has_page_field}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
