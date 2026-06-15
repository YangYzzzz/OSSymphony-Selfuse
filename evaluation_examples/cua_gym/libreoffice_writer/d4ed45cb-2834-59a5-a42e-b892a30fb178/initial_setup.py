"""
Initial Setup: Cut the last paragraph and paste it as the first paragraph.
Task ID: writer_edit_021
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_021'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/essay_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Paragraph 1
    p1 = doc.add_paragraph(
        "Sustainable development has emerged as one of the most pressing global priorities "
        "of the twenty-first century. As human populations grow and resource consumption "
        "intensifies, governments, businesses, and communities are increasingly compelled "
        "to rethink how economic progress can be achieved without undermining the natural "
        "systems that support all life on Earth."
    )

    # Paragraph 2
    p2 = doc.add_paragraph(
        "Environmental challenges such as climate change, biodiversity loss, and freshwater "
        "scarcity demand urgent action at every level of society. International agreements "
        "like the Paris Accord and the United Nations Sustainable Development Goals have "
        "provided frameworks for coordinated global responses, encouraging nations to adopt "
        "cleaner energy sources, protect critical ecosystems, and reduce greenhouse gas "
        "emissions across all sectors of their economies."
    )

    # Paragraph 3
    p3 = doc.add_paragraph(
        "Social equity is equally central to the concept of sustainability. Development "
        "strategies that ignore the needs of marginalized communities or perpetuate "
        "systemic inequalities ultimately fail to create lasting change. Inclusive policies "
        "that invest in education, healthcare, and economic opportunity for all people "
        "not only improve quality of life but also strengthen the social cohesion necessary "
        "for long-term stability and resilience."
    )

    # Paragraph 4 (last paragraph — this is the one the agent must move to first position)
    p4 = doc.add_paragraph(
        "In conclusion, sustainable development requires a balanced approach that considers "
        "economic growth, environmental protection, and social equity in equal measure."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
