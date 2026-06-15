"""
Reward Script: Two-column glossary layout with section break and paragraph formatting
Task ID: wrpara_046
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Continuous section break before Glossary heading
  Component 2 (0.30): Glossary section has 2 columns with ~0.5cm gap
  Component 3 (0.20): Glossary term paragraphs have 6pt space_before
  Component 4 (0.25): Glossary term paragraphs have 2cm hanging indent
"""

import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Emu

WORKDIR = '/home/user'
TASK_ID = 'wrpara_046'


def persist_app_state(domain):
    """Best-effort save via Ctrl+S in case the file is still open in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for %s" % domain)
    except Exception as e:
        print("PERSIST_WARN: save hook failed: %s" % e)


def find_glossary_heading_index(doc):
    """Find the index of the 'Glossary' heading paragraph."""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'glossary' and 'Heading' in (para.style.name or ''):
            return i
    # Fallback: look for any paragraph whose text is just 'Glossary'
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'glossary':
            return i
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    glossary_idx = find_glossary_heading_index(doc)
    if glossary_idx is None:
        print("CRITICAL: Could not find Glossary heading paragraph")
        print("REWARD: 0.0")
        return 0.0

    print("INFO: Glossary heading found at paragraph index %d" % glossary_idx)
    num_sections = len(doc.sections)
    print("INFO: Document has %d section(s)" % num_sections)

    # =========================================================================
    # Component 1: Continuous section break before Glossary heading (0.25 pts)
    # In OOXML, a section break is stored in the pPr of the LAST paragraph of
    # the PREVIOUS section. So the paragraph just before the Glossary heading
    # (index glossary_idx - 1) should contain a sectPr with type='continuous'.
    # =========================================================================
    try:
        if glossary_idx == 0:
            print("FAIL: Component 1 — Glossary heading is the first paragraph, no room for section break")
        else:
            prev_para = doc.paragraphs[glossary_idx - 1]
            pPr = prev_para._element.find(qn('w:pPr'))
            sectPr = None
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))

            # Derive booleans from actual XML query results
            sect_break_found = sectPr is not None
            type_el = sectPr.find(qn('w:type')) if sect_break_found else None
            is_continuous = (type_el is not None and type_el.get(qn('w:val')) == 'continuous')

            if sect_break_found and is_continuous:
                print("PASS: Component 1 — Continuous section break found before Glossary heading (0.25 pts)")
                total_score += 0.25
            elif sect_break_found:
                print("PARTIAL: Component 1 — Section break found but not continuous type (0.10 pts)")
                total_score += 0.10
            else:
                # Also check if there are multiple sections at all
                if num_sections >= 2:
                    print("PARTIAL: Component 1 — Multiple sections exist but break not in expected location (0.05 pts)")
                    total_score += 0.05
                else:
                    print("FAIL: Component 1 — No section break found before Glossary heading")
    except Exception as e:
        print("ERROR: Component 1 — %s" % e)

    # =========================================================================
    # Component 2: Glossary section has 2 columns with ~0.5cm gap (0.30 pts)
    # The glossary section is the LAST section (body-level sectPr).
    # Check: cols num=2, space ~283 twips (0.5cm = 283.5 twips).
    # =========================================================================
    try:
        if num_sections < 2:
            print("FAIL: Component 2 — Only 1 section, no separate glossary section")
        else:
            # The glossary section properties are in the body-level sectPr (last section)
            last_section = doc.sections[-1]
            sectPr = last_section._sectPr
            cols_el = sectPr.find(qn('w:cols'))

            col_num = None
            col_space = None
            if cols_el is not None:
                num_attr = cols_el.get(qn('w:num'))
                col_num = int(num_attr) if num_attr else None
                space_attr = cols_el.get(qn('w:space'))
                col_space = int(space_attr) if space_attr else None

            pts_earned = 0.0

            # Sub-check 2a: 2 columns (0.20 pts)
            if col_num == 2:
                print("PASS: Component 2a — Glossary section has 2 columns")
                pts_earned += 0.20
            else:
                print("FAIL: Component 2a — Expected 2 columns, found %s" % col_num)

            # Sub-check 2b: column gap ~0.5cm (283-284 twips) with tolerance (0.10 pts)
            # 0.5cm = 283.5 twips. Allow tolerance of +/- 15 twips (~0.026cm)
            if col_space is not None:
                gap_cm = col_space / 567.0
                if abs(col_space - 283.5) <= 15:
                    print("PASS: Component 2b — Column gap is %d twips (%.3f cm, target ~0.5cm)" % (col_space, gap_cm))
                    pts_earned += 0.10
                else:
                    print("FAIL: Component 2b — Column gap is %d twips (%.3f cm), expected ~283 twips (0.5cm)" % (col_space, gap_cm))
            else:
                print("FAIL: Component 2b — No column spacing attribute found")

            if pts_earned > 0:
                total_score += pts_earned
            print("INFO: Component 2 total — %.2f pts" % pts_earned)
    except Exception as e:
        print("ERROR: Component 2 — %s" % e)

    # =========================================================================
    # Component 3: Glossary term paragraphs have 6pt space_before (0.20 pts)
    # Glossary terms are paragraphs after the Glossary heading (indices glossary_idx+1 to end).
    # Each should have space_before = 76200 EMU (6pt).
    # =========================================================================
    try:
        glossary_paras = doc.paragraphs[glossary_idx + 1:]
        if not glossary_paras:
            print("FAIL: Component 3 — No glossary term paragraphs found after heading")
        else:
            target_sb_emu = 76200  # 6pt in EMU
            tolerance_emu = 2000  # small tolerance
            matching = 0
            total_terms = len(glossary_paras)

            for para in glossary_paras:
                sb = para.paragraph_format.space_before
                if sb is not None and abs(sb - target_sb_emu) <= tolerance_emu:
                    matching += 1

            ratio = matching / total_terms
            if ratio >= 0.9:
                print("PASS: Component 3 — %d/%d glossary paragraphs have 6pt space_before (0.20 pts)" % (matching, total_terms))
                total_score += 0.20
            elif ratio >= 0.5:
                partial = round(0.20 * ratio, 2)
                print("PARTIAL: Component 3 — %d/%d glossary paragraphs have 6pt space_before (%.2f pts)" % (matching, total_terms, partial))
                total_score += partial
            else:
                print("FAIL: Component 3 — Only %d/%d glossary paragraphs have 6pt space_before" % (matching, total_terms))
    except Exception as e:
        print("ERROR: Component 3 — %s" % e)

    # =========================================================================
    # Component 4: Glossary term paragraphs have 2cm hanging indent (0.25 pts)
    # Each glossary paragraph should have:
    #   left_indent ~= 720000 EMU (2cm)
    #   first_line_indent ~= -720000 EMU (-2cm, creating the hanging indent)
    # =========================================================================
    try:
        glossary_paras = doc.paragraphs[glossary_idx + 1:]
        if not glossary_paras:
            print("FAIL: Component 4 — No glossary term paragraphs found after heading")
        else:
            target_li_emu = 720000  # 2cm in EMU
            target_fli_emu = -720000  # -2cm in EMU
            tolerance_emu = 5000  # small tolerance (~0.014cm)
            matching_li = 0
            matching_fli = 0
            total_terms = len(glossary_paras)

            for para in glossary_paras:
                li = para.paragraph_format.left_indent
                fli = para.paragraph_format.first_line_indent

                if li is not None and abs(li - target_li_emu) <= tolerance_emu:
                    matching_li += 1
                if fli is not None and abs(fli - target_fli_emu) <= tolerance_emu:
                    matching_fli += 1

            li_ratio = matching_li / total_terms
            fli_ratio = matching_fli / total_terms

            pts_earned = 0.0

            # Sub-check 4a: left indent ~2cm (0.125 pts)
            if li_ratio >= 0.9:
                print("PASS: Component 4a — %d/%d paragraphs have 2cm left indent" % (matching_li, total_terms))
                pts_earned += 0.125
            elif li_ratio >= 0.5:
                partial = round(0.125 * li_ratio, 3)
                print("PARTIAL: Component 4a — %d/%d paragraphs have 2cm left indent (%.3f pts)" % (matching_li, total_terms, partial))
                pts_earned += partial
            else:
                print("FAIL: Component 4a — Only %d/%d paragraphs have 2cm left indent" % (matching_li, total_terms))

            # Sub-check 4b: first_line_indent ~-2cm (hanging) (0.125 pts)
            if fli_ratio >= 0.9:
                print("PASS: Component 4b — %d/%d paragraphs have -2cm first-line indent (hanging)" % (matching_fli, total_terms))
                pts_earned += 0.125
            elif fli_ratio >= 0.5:
                partial = round(0.125 * fli_ratio, 3)
                print("PARTIAL: Component 4b — %d/%d paragraphs have -2cm first-line indent (%.3f pts)" % (matching_fli, total_terms, partial))
                pts_earned += partial
            else:
                print("FAIL: Component 4b — Only %d/%d paragraphs have -2cm first-line indent" % (matching_fli, total_terms))

            if pts_earned > 0:
                total_score += pts_earned
            print("INFO: Component 4 total — %.3f pts" % pts_earned)
    except Exception as e:
        print("ERROR: Component 4 — %s" % e)

    final_score = min(round(total_score, 2), 1.0)
    print("")
    print("Score: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = '%s/%s.docx' % (WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
