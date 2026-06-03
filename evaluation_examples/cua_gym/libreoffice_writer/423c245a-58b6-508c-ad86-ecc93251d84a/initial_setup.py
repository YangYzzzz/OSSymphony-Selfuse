"""
Initial Setup: Technical manual with glossary section (single-column, unformatted glossary)
Task ID: wrpara_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_046'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('CloudStack Platform Technical Reference Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Version 3.2.1 | March 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    # --- Chapter 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'CloudStack Platform is an enterprise-grade infrastructure management solution '
        'designed for organizations requiring scalable, secure, and highly available '
        'computing environments. This manual covers installation procedures, architecture '
        'design principles, and operational guidelines for system administrators and '
        'DevOps engineers managing CloudStack deployments across hybrid cloud environments.'
    )
    doc.add_paragraph(
        'The platform supports multi-tenant resource isolation, automated failover '
        'mechanisms, and real-time monitoring dashboards. It integrates with popular '
        'identity providers such as Okta, Azure AD, and LDAP for centralized access '
        'management. CloudStack has been deployed in over 2,400 production environments '
        'worldwide since its initial release in 2019.'
    )

    # --- Chapter 2: System Architecture ---
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        'The CloudStack architecture follows a microservices-based design pattern with '
        'three primary layers: the Presentation Layer (web dashboard and CLI tools), '
        'the Service Layer (API gateway, orchestration engine, and task scheduler), and '
        'the Data Layer (distributed storage, relational databases, and caching systems).'
    )

    doc.add_heading('2.1 Presentation Layer', level=2)
    doc.add_paragraph(
        'The web dashboard is built using React 18 with TypeScript, communicating with '
        'backend services through a RESTful API gateway. The CLI tool, cstack-cli, is '
        'implemented in Go and supports batch operations, scriptable workflows, and '
        'SSH-based tunneling for secure remote administration. Both interfaces authenticate '
        'via OAuth 2.0 tokens issued by the integrated identity service.'
    )

    doc.add_heading('2.2 Service Layer', level=2)
    doc.add_paragraph(
        'The orchestration engine manages resource provisioning requests using a priority '
        'queue backed by RabbitMQ. Each provisioning request passes through validation, '
        'capacity checking, network allocation, and storage binding stages before final '
        'deployment. The average provisioning latency for a standard virtual machine is '
        '47 seconds under normal load conditions.'
    )

    doc.add_heading('2.3 Data Layer', level=2)
    doc.add_paragraph(
        'CloudStack uses PostgreSQL 15 for transactional data (user accounts, billing '
        'records, audit logs) and Apache Cassandra for time-series telemetry data. Object '
        'storage is provided through a MinIO-compatible interface supporting S3-compatible '
        'API calls. The system maintains three replicas of all critical data across '
        'geographically separated availability zones.'
    )

    # --- Chapter 3: Installation & Configuration ---
    doc.add_heading('3. Installation and Configuration', level=1)
    doc.add_paragraph(
        'Before installing CloudStack, ensure the target environment meets the minimum '
        'hardware requirements: 32 GB RAM, 8 CPU cores, and 500 GB SSD storage per management '
        'node. A minimum of three management nodes is recommended for production deployments '
        'to ensure high availability of the control plane.'
    )
    doc.add_paragraph(
        'The installation process involves downloading the CloudStack installer package '
        'from the official repository (https://packages.cloudstack.io/v3.2), verifying '
        'the GPG signature, and running the bootstrap script. The bootstrap process '
        'configures the local database, initializes the message broker, and registers '
        'the node with the cluster coordinator. Configuration is managed through YAML '
        'files located in /etc/cloudstack/conf.d/ on each management node.'
    )

    # --- Chapter 4: Networking ---
    doc.add_heading('4. Networking', level=1)
    doc.add_paragraph(
        'CloudStack supports VLAN-based isolation, VXLAN overlays, and software-defined '
        'networking (SDN) through integration with Open vSwitch. Each tenant receives an '
        'isolated virtual network with configurable subnet ranges, firewall rules, and '
        'load balancer policies. The platform supports both IPv4 and IPv6 addressing '
        'with dual-stack configurations available at the tenant level.'
    )
    doc.add_paragraph(
        'Network throughput is optimized through SR-IOV passthrough on supported hardware, '
        'achieving near-native performance with measured throughput of 9.4 Gbps on 10GbE '
        'interfaces. Quality of Service (QoS) policies can be applied per-tenant or '
        'per-virtual-machine to guarantee minimum bandwidth allocations during peak usage.'
    )

    # --- Chapter 5: Security ---
    doc.add_heading('5. Security', level=1)
    doc.add_paragraph(
        'All inter-service communication is encrypted using TLS 1.3 with mutual '
        'authentication. API tokens are rotated every 24 hours and stored in HashiCorp '
        'Vault. The platform supports role-based access control (RBAC) with predefined '
        'roles (Administrator, Operator, Viewer, Auditor) and custom role definitions '
        'using a JSON-based policy language.'
    )
    doc.add_paragraph(
        'Security audit logs capture all administrative actions with timestamps, source '
        'IP addresses, and affected resources. Logs are retained for 90 days in hot storage '
        'and archived to cold storage for seven years to comply with SOC 2 and ISO 27001 '
        'requirements. Intrusion detection is provided through integration with Falco and '
        'custom rule sets maintained by the CloudStack security team.'
    )

    # --- Chapter 6: Monitoring & Alerting ---
    doc.add_heading('6. Monitoring and Alerting', level=1)
    doc.add_paragraph(
        'The built-in monitoring subsystem collects metrics from all managed resources at '
        '15-second intervals. Metrics include CPU utilization, memory consumption, disk I/O, '
        'network traffic, and application-specific indicators exposed through a Prometheus-compatible '
        'endpoint. Alerting rules are defined using a threshold-based engine with support '
        'for multi-condition triggers and escalation policies.'
    )

    # --- Glossary --- (NO special formatting, single column, no section breaks)
    doc.add_heading('Glossary', level=1)

    glossary_terms = [
        ('API Gateway', 'A server that acts as the single entry point for all client '
         'requests, handling authentication, rate limiting, request routing, and response '
         'aggregation across backend microservices in the CloudStack platform.'),
        ('Availability Zone', 'A physically isolated data center location within a region '
         'that provides independent power, cooling, and network connectivity to ensure fault '
         'tolerance and high availability of deployed workloads.'),
        ('CIDR Block', 'Classless Inter-Domain Routing notation used to specify IP address '
         'ranges for virtual network subnets, security group rules, and access control lists '
         'within the CloudStack networking subsystem.'),
        ('Container Orchestration', 'The automated management of containerized application '
         'lifecycles including deployment, scaling, networking, and health monitoring, '
         'implemented through integration with Kubernetes 1.28 in CloudStack.'),
        ('Data Replication', 'The process of maintaining synchronized copies of data across '
         'multiple storage nodes or availability zones to ensure durability and provide '
         'read scaling for high-throughput workloads.'),
        ('Edge Computing', 'A distributed computing paradigm that brings computation and '
         'data storage closer to end users, reducing latency and bandwidth consumption '
         'for geographically dispersed CloudStack deployments.'),
        ('Failover', 'The automatic transfer of workloads from a failed or degraded node '
         'to a healthy standby node, ensuring continuous service availability with minimal '
         'interruption as defined by the configured recovery time objective.'),
        ('Horizontal Scaling', 'The practice of adding additional compute instances to '
         'distribute workload across multiple nodes, as opposed to vertical scaling which '
         'increases the resources of existing instances in the cluster.'),
        ('Idempotency', 'A property of API operations ensuring that multiple identical '
         'requests produce the same result as a single request, critical for reliable '
         'retry mechanisms in distributed systems like CloudStack.'),
        ('Key Management Service', 'A centralized system for creating, storing, rotating, '
         'and revoking cryptographic keys used for data encryption, TLS certificates, and '
         'API token signing across all CloudStack components.'),
        ('Load Balancer', 'A component that distributes incoming network traffic across '
         'multiple backend servers using configurable algorithms such as round-robin, '
         'least-connections, or weighted response time balancing.'),
        ('Namespace', 'A logical isolation boundary within CloudStack that partitions '
         'resources, configurations, and access policies to support multi-tenant environments '
         'without interference between different organizational units.'),
        ('Object Storage', 'A flat-address storage architecture that manages data as objects '
         'with associated metadata and unique identifiers, providing an S3-compatible API '
         'for unstructured data such as backups, logs, and media files.'),
        ('Rate Limiting', 'A traffic management technique that restricts the number of '
         'API requests a client can make within a specified time window, protecting backend '
         'services from overload and ensuring fair resource allocation.'),
        ('Service Mesh', 'An infrastructure layer that manages service-to-service '
         'communication within the CloudStack microservices architecture, providing traffic '
         'management, observability, and security features through sidecar proxies.'),
    ]

    for term, definition in glossary_terms:
        para = doc.add_paragraph(f'{term} \u2013 {definition}')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
