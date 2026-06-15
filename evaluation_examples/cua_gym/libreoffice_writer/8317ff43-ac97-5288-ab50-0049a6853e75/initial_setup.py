"""
Initial Setup: NDA with tracked deletion in Section 4
Task ID: writer_rm_004
Domain: libreoffice_writer

Creates a Mutual NDA document with multiple sections.
In Section 4, a tracked deletion by 'James Rivera' removes the confidentiality
clause paragraph. There are also other tracked changes elsewhere in the document.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_tracked_deletion(paragraph, text, author, date="2026-03-28T14:22:00Z", run_id="1"):
    """Add a tracked deletion (w:del) to a paragraph using raw XML."""
    # Create the w:del element wrapping a w:r with w:delText
    del_elem = parse_xml(
        f'<w:del {nsdecls("w")} w:id="{run_id}" w:author="{author}" w:date="{date}">'
        f'  <w:r>'
        f'    <w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'    <w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>'
        f'    <w:delText xml:space="preserve">{text}</w:delText>'
        f'  </w:r>'
        f'</w:del>'
    )
    paragraph._element.append(del_elem)


def add_tracked_insertion(paragraph, text, author, date, run_id="10",
                          bold=False, font_name=None, font_size=None):
    """Add a tracked insertion (w:ins) to a paragraph using raw XML."""
    rpr_parts = []
    if bold:
        rpr_parts.append('<w:b/>')
    if font_name:
        rpr_parts.append(f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
    if font_size:
        val = str(int(font_size * 2))  # half-points
        rpr_parts.append(f'<w:sz w:val="{val}"/><w:szCs w:val="{val}"/>')
    rpr_xml = '<w:rPr>' + ''.join(rpr_parts) + '</w:rPr>' if rpr_parts else ''

    ins_elem = parse_xml(
        f'<w:ins {nsdecls("w")} w:id="{run_id}" w:author="{author}" w:date="{date}">'
        f'  <w:r>{rpr_xml}<w:t xml:space="preserve">{text}</w:t></w:r>'
        f'</w:ins>'
    )
    paragraph._element.append(ins_elem)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- Title ---
    title = doc.add_heading('MUTUAL NON-DISCLOSURE AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Preamble ---
    doc.add_paragraph(
        'This Mutual Non-Disclosure Agreement ("Agreement") is entered into as of '
        'March 15, 2026, by and between Vertex Analytics Inc., a Delaware corporation '
        '("Company A"), and Meridian Data Solutions LLC, a California limited liability '
        'company ("Company B").'
    )
    doc.add_paragraph(
        'WHEREAS, the parties wish to explore a potential business relationship '
        '("Purpose") and, in connection therewith, may disclose to each other certain '
        'confidential and proprietary information; and'
    )
    doc.add_paragraph(
        'WHEREAS, the parties desire to protect such information from unauthorized '
        'use and disclosure;'
    )
    doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements '
        'set forth herein, the parties agree as follows:'
    )

    # --- Section 1: Definitions ---
    doc.add_heading('Section 1: Definitions', level=1)
    doc.add_paragraph(
        '"Confidential Information" means any and all non-public, proprietary, or '
        'confidential information disclosed by either party to the other, whether orally, '
        'in writing, electronically, or by inspection of tangible objects. This includes, '
        'but is not limited to: trade secrets, business plans, financial data, customer '
        'lists, technical specifications, source code, algorithms, marketing strategies, '
        'product roadmaps, and personnel information.'
    )
    doc.add_paragraph(
        '"Disclosing Party" means the party disclosing Confidential Information.'
    )
    doc.add_paragraph(
        '"Receiving Party" means the party receiving Confidential Information.'
    )

    # --- Section 2: Scope of Agreement ---
    doc.add_heading('Section 2: Scope of Agreement', level=1)
    doc.add_paragraph(
        'This Agreement shall govern all Confidential Information exchanged between '
        'the parties from the Effective Date through the termination of this Agreement. '
        'The obligations herein shall apply to all officers, directors, employees, agents, '
        'and contractors of each party who may have access to Confidential Information.'
    )

    # Add a tracked insertion in Section 2 (by a different author)
    p_insert = doc.add_paragraph()
    add_tracked_insertion(
        p_insert,
        'The scope of this Agreement extends to any subsidiaries or affiliates '
        'of either party that may be involved in the Purpose.',
        author='Lisa Thompson',
        date='2026-03-27T10:15:00Z',
        run_id='11',
        font_name='Times New Roman',
        font_size=12
    )

    # --- Section 3: Exclusions ---
    doc.add_heading('Section 3: Exclusions', level=1)
    doc.add_paragraph(
        'Confidential Information shall not include information that: (a) is or becomes '
        'publicly available through no fault of the Receiving Party; (b) was already known '
        'to the Receiving Party prior to disclosure; (c) is independently developed by '
        'the Receiving Party without use of or reference to the Confidential Information; '
        'or (d) is rightfully received from a third party without restriction on disclosure.'
    )

    # --- Section 4: Confidentiality Obligations ---
    doc.add_heading('Section 4: Confidentiality Obligations', level=1)
    doc.add_paragraph(
        'The Receiving Party shall use the Confidential Information solely for the Purpose '
        'and shall not disclose it to any third party without the prior written consent of '
        'the Disclosing Party.'
    )

    # THIS IS THE KEY TRACKED DELETION by James Rivera
    # The paragraph is shown as a tracked deletion — the entire paragraph text is inside w:del
    del_para = doc.add_paragraph()
    add_tracked_deletion(
        del_para,
        'Both parties agree to maintain strict confidentiality of all shared information '
        'and shall implement reasonable security measures, including but not limited to '
        'encryption of digital files, restricted access controls, secure storage of physical '
        'documents, and mandatory confidentiality training for all personnel with access '
        'to the disclosed materials.',
        author='James Rivera',
        date='2026-03-28T14:22:00Z',
        run_id='2'
    )

    doc.add_paragraph(
        'The Receiving Party shall limit access to Confidential Information to those '
        'employees and contractors who have a need to know and who are bound by '
        'confidentiality obligations no less restrictive than those set forth herein.'
    )

    # --- Section 5: Term and Termination ---
    doc.add_heading('Section 5: Term and Termination', level=1)

    # Another tracked change: insertion by Lisa Thompson
    p5 = doc.add_paragraph(
        'This Agreement shall remain in effect for a period of '
    )
    add_tracked_insertion(
        p5,
        'three (3)',
        author='Lisa Thompson',
        date='2026-03-27T09:45:00Z',
        run_id='12',
        font_name='Times New Roman',
        font_size=12
    )
    run_after = p5.add_run(
        ' years from the Effective Date, unless terminated earlier by either party '
        'upon thirty (30) days written notice to the other party.'
    )

    doc.add_paragraph(
        'The obligations of confidentiality shall survive termination of this Agreement '
        'for a period of five (5) years following such termination.'
    )

    # --- Section 6: Remedies ---
    doc.add_heading('Section 6: Remedies', level=1)
    doc.add_paragraph(
        'Each party acknowledges that a breach of this Agreement may cause irreparable '
        'harm to the other party for which monetary damages would be inadequate. '
        'Accordingly, each party shall be entitled to seek equitable relief, including '
        'injunction and specific performance, in addition to all other remedies available '
        'at law or in equity.'
    )

    # A tracked format change in Section 6 (bold insertion)
    p6_track = doc.add_paragraph()
    add_tracked_insertion(
        p6_track,
        'Liquidated damages of $50,000 per breach shall apply where actual damages '
        'cannot be readily ascertained.',
        author='James Rivera',
        date='2026-03-28T15:00:00Z',
        run_id='13',
        bold=True,
        font_name='Times New Roman',
        font_size=12
    )

    # --- Section 7: Governing Law ---
    doc.add_heading('Section 7: Governing Law', level=1)
    doc.add_paragraph(
        'This Agreement shall be governed by and construed in accordance with the laws '
        'of the State of Delaware, without regard to its conflict of laws principles. '
        'Any disputes arising under this Agreement shall be resolved in the state or '
        'federal courts located in Wilmington, Delaware.'
    )

    # --- Signature Block ---
    doc.add_paragraph('')  # spacer
    sig = doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement '
                            'as of the date first written above.')
    doc.add_paragraph('')

    # Signature lines
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('Vertex Analytics Inc.')
    doc.add_paragraph('By: Elena Kowalski, Chief Executive Officer')
    doc.add_paragraph('Date: March 15, 2026')
    doc.add_paragraph('')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('Meridian Data Solutions LLC')
    doc.add_paragraph('By: David Park, Managing Director')
    doc.add_paragraph('Date: March 15, 2026')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
