"""
Initial Setup: Research Paper Template document with custom styles
Task ID: writer_bs_063
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

def create_initial():
    doc = Document()

    # --- Page Setup: 2.54cm margins ---
    section = doc.sections[0]
    margin = Cm(2.54)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)  # A4

    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Research Paper Draft"
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in hp.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Footer with page number ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fp.add_run("Page ")
    r1 = fp.add_run()
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))
    for run in fp.runs:
        run.font.size = Pt(9)

    # --- Custom Styles ---
    styles = doc.styles

    # Heading 1 style customization
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2 style customization
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(15)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0x2B, 0x57, 0x97)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Heading 3 style customization
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Arial'
    h3_font.size = Pt(13)
    h3_font.bold = True
    h3_font.italic = True
    h3_font.color.rgb = RGBColor(0x3A, 0x6E, 0xA5)
    h3_style.paragraph_format.space_before = Pt(14)
    h3_style.paragraph_format.space_after = Pt(6)

    # Body Text style customization
    bt_style = styles['Body Text']
    bt_font = bt_style.font
    bt_font.name = 'Times New Roman'
    bt_font.size = Pt(12)
    bt_style.paragraph_format.line_spacing = 1.5
    bt_style.paragraph_format.space_after = Pt(6)

    # --- Document Content ---
    # Title
    title = doc.add_heading('The Impact of Machine Learning on Modern Healthcare Systems', level=1)

    # Abstract section
    doc.add_heading('Abstract', level=2)
    abstract = doc.add_paragraph(style='Body Text')
    abstract.text = (
        'This paper examines the transformative role of machine learning algorithms '
        'in contemporary healthcare delivery. Through a comprehensive analysis of '
        'recent advancements in diagnostic imaging, drug discovery, and patient '
        'outcome prediction, we demonstrate that ML-based approaches consistently '
        'outperform traditional statistical methods across multiple clinical domains. '
        'Our findings suggest that the integration of these technologies into clinical '
        'workflows can reduce diagnostic errors by up to 34% while simultaneously '
        'decreasing average patient wait times.'
    )

    # Introduction
    doc.add_heading('1. Introduction', level=2)
    intro1 = doc.add_paragraph(style='Body Text')
    intro1.text = (
        'The healthcare industry generates approximately 2.3 exabytes of data annually, '
        'encompassing electronic health records, medical imaging, genomic sequencing, '
        'and wearable device telemetry. Traditional analytical approaches have proven '
        'insufficient for extracting actionable insights from datasets of this magnitude '
        'and complexity (Roberts et al., 2024).'
    )
    intro2 = doc.add_paragraph(style='Body Text')
    intro2.text = (
        'Machine learning, a subset of artificial intelligence that enables systems to '
        'learn from data without explicit programming, has emerged as a promising '
        'framework for addressing these challenges. Recent breakthroughs in deep learning '
        'architectures, particularly convolutional neural networks and transformer models, '
        'have demonstrated remarkable accuracy in tasks ranging from tumor detection to '
        'protein structure prediction (Chen & Williams, 2025).'
    )

    # Methodology
    doc.add_heading('2. Methodology', level=2)

    doc.add_heading('2.1 Data Collection', level=3)
    method1 = doc.add_paragraph(style='Body Text')
    method1.text = (
        'We conducted a systematic review of 847 peer-reviewed publications from '
        'January 2020 through December 2025, sourced from PubMed, IEEE Xplore, and '
        'the ACM Digital Library. Inclusion criteria required empirical validation on '
        'clinical datasets with a minimum of 500 patient records.'
    )

    doc.add_heading('2.2 Analysis Framework', level=3)
    method2 = doc.add_paragraph(style='Body Text')
    method2.text = (
        'Each study was evaluated against five performance metrics: sensitivity, '
        'specificity, area under the ROC curve (AUC), positive predictive value, '
        'and clinical workflow integration score. The latter metric, developed by '
        'our team, quantifies the practical feasibility of deploying the proposed '
        'solution within existing hospital information systems.'
    )

    # Results
    doc.add_heading('3. Results', level=2)
    results1 = doc.add_paragraph(style='Body Text')
    results1.text = (
        'Among the reviewed studies, 612 (72.3%) reported statistically significant '
        'improvements over baseline methods. The median AUC across all ML-based '
        'diagnostic tools was 0.923 (IQR: 0.891-0.956), compared to 0.847 '
        '(IQR: 0.812-0.883) for conventional approaches. Particularly notable gains '
        'were observed in radiology (mean improvement: 14.2%), pathology (11.7%), '
        'and dermatology (9.8%).'
    )

    # References
    doc.add_heading('References', level=2)
    refs = [
        'Chen, L. & Williams, R. (2025). Transformer architectures for clinical decision support. Nature Medicine, 31(4), 445-458.',
        'Roberts, A., Patel, S., & Kim, J. (2024). Big data challenges in modern healthcare systems. Journal of Health Informatics, 18(2), 112-129.',
        'Thompson, D. et al. (2023). Deep learning in diagnostic radiology: A systematic review. Radiology, 308(3), e230142.',
    ]
    for ref in refs:
        p = doc.add_paragraph(style='Body Text')
        p.text = ref
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)  # hanging indent

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')

create_initial()
