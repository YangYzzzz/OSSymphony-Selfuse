"""
Initial Setup: Word sorting exercise - color-coded by word type
Task ID: osworld_writer_colorword_009
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_009'
OUTPUT = f'{WORKDIR}/word_sorting_exercise.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Word Sorting Exercise")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Add instruction paragraph
    instr_para = doc.add_paragraph(
        "Sort the following words by type: action words (verbs) and descriptive words (adjectives)."
    )

    # 54 words arranged in a 9x6 table
    # Includes specific words from context:
    #   'runs', 'jumped', 'singing' (verbs), 'careful', 'homeless', 'quickly' (adjectives), 'tree' (other)
    # Additional realistic mix of verbs, adjectives, and other words
    words = [
        # Row 1
        "runs", "careful", "tree", "jumped", "quickly", "river",
        # Row 2
        "singing", "homeless", "mountain", "builds", "peaceful", "stone",
        # Row 3
        "walked", "joyful", "cloud", "reads", "helpless", "book",
        # Row 4
        "writes", "graceful", "bridge", "talks", "endless", "forest",
        # Row 5
        "plays", "colorful", "wind", "jumps", "restless", "garden",
        # Row 6
        "dancing", "wonderful", "lamp", "drawing", "brightly", "chair",
        # Row 7
        "fixes", "thankful", "pencil", "paints", "softly", "window",
        # Row 8
        "prints", "fearless", "mirror", "lifts", "loudly", "table",
        # Row 9
        "checks", "skillful", "bottle", "grows", "slowly", "paper",
    ]

    # Verify we have exactly 54 words
    assert len(words) == 54, f"Expected 54 words, got {len(words)}"

    # Create 9x6 table
    table = doc.add_table(rows=9, cols=6)
    table.style = 'Table Grid'

    for row_idx in range(9):
        for col_idx in range(6):
            word_idx = row_idx * 6 + col_idx
            cell = table.cell(row_idx, col_idx)
            # Clear any existing content
            for para in cell.paragraphs:
                for run in para.runs:
                    run.clear()

            para = cell.paragraphs[0]
            run = para.add_run(words[word_idx])
            # All words in default black text — NO coloring applied
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
