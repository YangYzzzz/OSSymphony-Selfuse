"""
Reward Script: Color-code words in a sorting exercise table
Task ID: osworld_writer_colorword_009
Domain: libreoffice_writer
Scoring:
  - Component 1: Orange coloring for verbs (-s, -ed, -ing endings) — 0.5 pts
  - Component 2: Teal coloring for adjectives (-ful, -less, -ly endings) — 0.5 pts
  Note: Neutral words remaining black is a gate check (precondition), not scored separately
  because it is true in both initial and golden states.
  Total: 1.0
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_009'

# Target colors
# Orange: FFA500
ORANGE = (0xFF, 0xA5, 0x00)
# Teal: 008080
TEAL = (0x00, 0x80, 0x80)
# Black: 000000
BLACK = (0x00, 0x00, 0x00)

# Color tolerance for perceptual comparison
COLOR_TOLERANCE = 30


def color_distance(c1, c2):
    """Euclidean distance between two RGB colors."""
    return sqrt(sum((a - b) ** 2 for a, b in zip(c1, c2)))


def is_orange(rgb):
    """Check if color is close to orange (FFA500)."""
    return color_distance(rgb, ORANGE) < COLOR_TOLERANCE


def is_teal(rgb):
    """Check if color is close to teal (008080)."""
    return color_distance(rgb, TEAL) < COLOR_TOLERANCE


def is_black(rgb):
    """Check if color is close to black (000000)."""
    return color_distance(rgb, BLACK) < COLOR_TOLERANCE


def get_word_color(cell):
    """
    Get the effective font color of a table cell.
    Returns (r, g, b) tuple or None if no explicit color found.
    """
    for para in cell.paragraphs:
        for run in para.runs:
            try:
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    return (rgb[0], rgb[1], rgb[2])
            except Exception:
                pass
    return None


def classify_word(word):
    """
    Classify a word based on its ending:
    - 'adj'    → ends with -ful, -less, or -ly → should be teal (checked FIRST)
    - 'verb'   → ends with -s, -ed, or -ing → should be orange
    - 'neutral'→ all others → should be black

    IMPORTANT: Adjective suffixes must be checked before verb suffixes because
    words like 'homeless', 'restless', 'fearless' end in '-less' (adj) but also
    end in '-s' (verb). The more specific adjective check takes priority.
    """
    w = word.lower().strip()
    # Check adjective endings FIRST (more specific patterns)
    if w.endswith('ful') or w.endswith('less') or w.endswith('ly'):
        return 'adj'
    # Then check verb endings
    elif w.endswith('ing') or w.endswith('ed') or w.endswith('s'):
        return 'verb'
    else:
        return 'neutral'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Verify document has exactly 1 table with 9x6 layout (precondition gate)
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    if num_rows != 9 or num_cols != 6:
        print(f"CRITICAL: Expected 9x6 table, found {num_rows}x{num_cols}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all words and their colors
    verb_results = []    # list of (word, actual_color)
    adj_results = []     # list of (word, actual_color)
    neutral_results = [] # list of (word, actual_color)

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            word = cell.text.strip()
            if not word:
                continue

            word_class = classify_word(word)
            actual_color = get_word_color(cell)

            if word_class == 'verb':
                verb_results.append((word, actual_color))
            elif word_class == 'adj':
                adj_results.append((word, actual_color))
            else:
                neutral_results.append((word, actual_color))

    print(f"Found: {len(verb_results)} verb words, {len(adj_results)} adj words, {len(neutral_results)} neutral words")

    # Component 1: Verb words (ending in -s, -ed, -ing) should be colored orange (FFA500) — 0.5 pts
    # This FAILS on initial_env (all black) and PASSES on golden_env (verbs are orange)
    try:
        verb_correct = 0
        verb_total = len(verb_results)

        for word, actual_color in verb_results:
            if actual_color is not None and is_orange(actual_color):
                verb_correct += 1
            else:
                color_str = f"rgb{actual_color}" if actual_color else "None/default(black)"
                print(f"  FAIL verb: '{word}' color={color_str} (expected orange FFA500)")

        if verb_total > 0 and verb_correct == verb_total:
            print(f"PASS: Component 1 — All {verb_total} verb words colored orange FFA500 (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — {verb_correct}/{verb_total} verb words correctly orange")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Adjective words (ending in -ful, -less, -ly) should be colored teal (008080) — 0.5 pts
    # This FAILS on initial_env (all black) and PASSES on golden_env (adjectives are teal)
    try:
        adj_correct = 0
        adj_total = len(adj_results)

        for word, actual_color in adj_results:
            if actual_color is not None and is_teal(actual_color):
                adj_correct += 1
            else:
                color_str = f"rgb{actual_color}" if actual_color else "None/default(black)"
                print(f"  FAIL adj: '{word}' color={color_str} (expected teal 008080)")

        if adj_total > 0 and adj_correct == adj_total:
            print(f"PASS: Component 2 — All {adj_total} adjective words colored teal 008080 (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 2 — {adj_correct}/{adj_total} adjective words correctly teal")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Gate check (no points): Neutral words should remain black
    # This is true in BOTH initial and golden, so it is NOT a scoring component.
    # It is logged for diagnostic purposes only.
    try:
        neutral_correct = 0
        neutral_total = len(neutral_results)
        for word, actual_color in neutral_results:
            if actual_color is None or is_black(actual_color):
                neutral_correct += 1
            else:
                color_str = f"rgb{actual_color}"
                print(f"  WARN neutral: '{word}' color={color_str} (expected black 000000)")
        print(f"INFO Gate: {neutral_correct}/{neutral_total} neutral words are black (no points — precondition)")
    except Exception as e:
        print(f"ERROR: Gate check — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/word_sorting_exercise.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
