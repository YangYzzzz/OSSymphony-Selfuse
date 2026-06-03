"""
Reward Script: Yard Sale Price List - Convert plain text to formatted table
Task ID: writer_creative_048
Domain: libreoffice_writer
Scoring:
  Component 1: Title formatting (centered, 18pt, bold)         — 0.25 points
  Component 2: Table structure (2 cols, 16 rows, correct headers) — 0.25 points
  Component 3: Header row formatting (bold + background #D9E2F3) — 0.25 points
  Component 4: Price column right-aligned with $ formatting    — 0.25 points
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_048'
FILE_PATH = f'{WORKDIR}/Desktop/yard_sale_prices.docx'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Initial state: Plain text paragraphs, no table, no formatted title.
    Golden state: Centered 18pt bold title + 2-column table with 16 rows,
                  header row in blue-gray background, price column right-aligned.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Title formatting — 0.25 points
    # The title 'Yard Sale Price List' should be centered, 18pt, and bold.
    # In initial_env it is left-aligned, 12pt, not bold — so this FAILS initially.
    # -----------------------------------------------------------------------
    try:
        title_para = None
        for para in doc.paragraphs:
            if 'Yard Sale Price List' in para.text:
                title_para = para
                break

        if title_para is None:
            print("FAIL: Component 1 — Title paragraph 'Yard Sale Price List' not found")
        else:
            alignment = title_para.paragraph_format.alignment
            is_centered = (alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

            # Check font size and bold across all runs
            runs = [r for r in title_para.runs if r.text.strip()]
            is_bold = any(r.bold for r in runs) if runs else False
            font_size_ok = False
            for r in runs:
                if r.font.size and abs(r.font.size.pt - 18.0) < 0.5:
                    font_size_ok = True
                    break

            if is_centered and is_bold and font_size_ok:
                print(f"PASS: Component 1 — Title is centered={is_centered}, bold={is_bold}, 18pt={font_size_ok} (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — Title centered={is_centered}, bold={is_bold}, 18pt={font_size_ok} (expected all True)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Table structure — 0.25 points
    # A 2-column table with 16 rows (1 header + 15 data rows) and correct headers.
    # In initial_env there is no table at all — so this FAILS initially.
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No table found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)

            header_row = table.rows[0]
            header_item = header_row.cells[0].text.strip()
            header_price = header_row.cells[1].text.strip()

            structure_ok = (num_rows == 16 and num_cols == 2)
            headers_ok = (header_item.lower() == 'item' and header_price.lower() == 'price')

            if structure_ok and headers_ok:
                print(f"PASS: Component 2 — Table has {num_rows} rows x {num_cols} cols, headers='{header_item}'/'{header_price}' (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — rows={num_rows} (expected 16), cols={num_cols} (expected 2), headers='{header_item}'/'{header_price}' (expected 'Item'/'Price')")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Header row formatting — 0.25 points
    # Header row cells should have bold text and background color #D9E2F3.
    # In initial_env there is no table — so this FAILS initially.
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No table found")
        else:
            table = doc.tables[0]
            header_row = table.rows[0]

            # Check background color of header cells
            target_color = 'D9E2F3'
            bg_ok = True
            for cell in header_row.cells:
                tc = cell._tc
                shading = tc.find('.//' + qn('w:shd'))
                fill = shading.get(qn('w:fill')) if shading is not None else None
                if fill is None or fill.upper() != target_color.upper():
                    bg_ok = False
                    break

            # Check bold in header cells
            bold_ok = True
            for cell in header_row.cells:
                cell_runs = [r for r in cell.paragraphs[0].runs if r.text.strip()]
                if not cell_runs:
                    # Check if the cell text exists but with no runs (possible style inheritance)
                    # Fall back: consider it bold if cell text is non-empty
                    pass
                else:
                    if not any(r.bold for r in cell_runs):
                        bold_ok = False
                        break

            if bg_ok and bold_ok:
                print(f"PASS: Component 3 — Header row has background #D9E2F3 and bold text (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — Header bg_ok={bg_ok}, bold_ok={bold_ok} (expected both True)")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Price column right-aligned with $ formatting — 0.25 points
    # All price column cells (rows 1-15) should be right-aligned and
    # contain values formatted as dollar amounts (e.g. '$35', '$8').
    # In initial_env there is no table — so this FAILS initially.
    # -----------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 4 — No table found")
        else:
            table = doc.tables[0]
            data_rows = table.rows[1:]  # skip header row

            right_aligned_count = 0
            dollar_format_count = 0
            total_data_rows = len(data_rows)

            for row in data_rows:
                price_cell = row.cells[1]
                price_text = price_cell.text.strip()
                price_align = price_cell.paragraphs[0].paragraph_format.alignment

                if price_align == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    right_aligned_count += 1

                if price_text.startswith('$') and len(price_text) > 1:
                    dollar_format_count += 1

            # Require all 15 data rows to pass both checks
            all_right_aligned = (right_aligned_count == total_data_rows)
            all_dollar_format = (dollar_format_count == total_data_rows)

            if all_right_aligned and all_dollar_format:
                print(f"PASS: Component 4 — All {total_data_rows} price cells right-aligned with $ format (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 4 — right_aligned={right_aligned_count}/{total_data_rows}, dollar_format={dollar_format_count}/{total_data_rows}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
