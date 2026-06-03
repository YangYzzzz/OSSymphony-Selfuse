"""
Initial Setup: Yard Sale Price List - plain text version
Task ID: writer_creative_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'yard_sale_prices'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Title paragraph - plain, 12pt, left-aligned (NOT formatted yet)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Yard Sale Price List')
    title_run.font.size = Pt(12)
    title_run.bold = False
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 15 items as plain text lines, 12pt, left-aligned
    items = [
        'Wooden bookshelf - $35',
        'Desk lamp - $8',
        'Set of 4 dinner plates - $12',
        'Vintage picture frame - $5',
        'Kids bicycle (16 inch) - $25',
        'Toaster - $10',
        'Yoga mat - $7',
        'Board game collection (5 games) - $15',
        'Table fan - $12',
        'Gardening tool set - $18',
        'DVD player - $10',
        'Winter coat (women\u2019s M) - $20',
        'Coffee maker - $15',
        'Throw pillows (set of 3) - $8',
        'Camping chair - $12',
    ]

    for item_text in items:
        para = doc.add_paragraph()
        run = para.add_run(item_text)
        run.font.size = Pt(12)
        run.bold = False
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
