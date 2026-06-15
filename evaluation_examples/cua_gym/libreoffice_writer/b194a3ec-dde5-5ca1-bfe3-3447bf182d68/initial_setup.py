"""
Initial Setup: Set alt text of image on page 1 to 'Company Logo - Acme Corp'
Task ID: writer_obj_015
Domain: libreoffice_writer

Creates accessible_doc.docx at /home/user/Desktop/ with:
- Page 1: Company letterhead with a 4cm x 4cm logo image (NO alt text set)
- Realistic company document content
- Image has NO alternative text (that is the task to add)
"""

import os
import io
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_015'
OUTPUT = f'{WORKDIR}/accessible_doc.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_logo_image() -> bytes:
    """Create a simple company logo image (4cm x 4cm at 96dpi ~ 151x151px)."""
    size = (160, 160)
    img = Image.new('RGB', size, color=(41, 128, 185))  # Corporate blue
    draw = ImageDraw.Draw(img)

    # Draw a simple logo shape - circle with letter
    draw.ellipse([10, 10, 150, 150], fill=(255, 255, 255), outline=(41, 128, 185), width=3)
    draw.ellipse([30, 30, 130, 130], fill=(41, 128, 185), outline=(255, 255, 255), width=2)

    # Draw 'A' for Acme
    draw.polygon([(80, 40), (50, 120), (110, 120)], outline=(255, 255, 255), fill=None)
    draw.line([(60, 95), (100, 95)], fill=(255, 255, 255), width=4)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def create_initial():
    """Create the initial accessible_doc.docx with image but NO alt text."""
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Page 1: Company letterhead with logo ---

    # Add the logo image (4cm x 4cm) - NO alt text set (that's the task)
    logo_buf = create_logo_image()
    logo_para = doc.add_paragraph()
    logo_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_buf, width=Cm(4), height=Cm(4))
    # NOTE: No alt text (description) is set on this image - that is the task

    # Company name heading
    heading = doc.add_paragraph()
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading.paragraph_format.space_before = Pt(6)
    heading_run = heading.add_run('Acme Corp')
    heading_run.font.name = 'Calibri'
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(41, 128, 185)

    # Tagline
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('Building Tomorrow, Today')
    tagline_run.font.name = 'Calibri'
    tagline_run.font.size = Pt(11)
    tagline_run.italic = True
    tagline_run.font.color.rgb = RGBColor(100, 100, 100)

    # Separator line (using border-like formatting via empty paragraph)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(12)

    # Document title
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run('Annual Accessibility Compliance Report')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    # Date and reference
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_after = Pt(18)
    date_run = date_para.add_run('Fiscal Year 2024  |  Document Ref: ACC-2024-015')
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(120, 120, 120)

    # Section 1: Executive Summary
    s1_heading = doc.add_paragraph()
    s1_heading.paragraph_format.space_before = Pt(12)
    s1_heading.paragraph_format.space_after = Pt(6)
    s1_run = s1_heading.add_run('1. Executive Summary')
    s1_run.font.name = 'Calibri'
    s1_run.font.size = Pt(13)
    s1_run.font.bold = True
    s1_run.font.color.rgb = RGBColor(41, 128, 185)

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'This report outlines Acme Corp\'s commitment to digital accessibility standards '
        'in accordance with WCAG 2.1 guidelines and Section 508 compliance requirements. '
        'Our accessibility program has achieved significant milestones during fiscal year 2024, '
        'improving user experience for over 12,000 employees and 45,000 external customers.'
    )
    intro_run.font.name = 'Calibri'
    intro_run.font.size = Pt(11)
    intro.paragraph_format.space_after = Pt(8)

    # Section 2: Key Findings
    s2_heading = doc.add_paragraph()
    s2_heading.paragraph_format.space_before = Pt(10)
    s2_heading.paragraph_format.space_after = Pt(6)
    s2_run = s2_heading.add_run('2. Key Findings')
    s2_run.font.name = 'Calibri'
    s2_run.font.size = Pt(13)
    s2_run.font.bold = True
    s2_run.font.color.rgb = RGBColor(41, 128, 185)

    findings = [
        'Image alternative text coverage increased from 34% to 78% across all web properties.',
        'Screen reader compatibility improved on 23 critical customer-facing applications.',
        'Color contrast ratios now meet AA standards for 91% of UI components.',
        'Keyboard navigation support added to all primary workflow interfaces.',
        'Training program delivered to 847 content creators and developers.',
    ]

    for finding in findings:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet_run = bullet.add_run(finding)
        bullet_run.font.name = 'Calibri'
        bullet_run.font.size = Pt(11)

    # Section 3: Accessibility Metrics Table
    s3_heading = doc.add_paragraph()
    s3_heading.paragraph_format.space_before = Pt(14)
    s3_heading.paragraph_format.space_after = Pt(6)
    s3_run = s3_heading.add_run('3. Accessibility Metrics Overview')
    s3_run.font.name = 'Calibri'
    s3_run.font.size = Pt(13)
    s3_run.font.bold = True
    s3_run.font.color.rgb = RGBColor(41, 128, 185)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Category', 'FY2023 Score', 'FY2024 Score', 'Status']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell_run = cell.paragraphs[0].add_run(header)
        cell_run.font.bold = True
        cell_run.font.name = 'Calibri'
        cell_run.font.size = Pt(10)
        cell_run.font.color.rgb = RGBColor(255, 255, 255)
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2980B9')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    data_rows = [
        ['Image Alt Text', '34%', '78%', 'Improving'],
        ['Color Contrast', '67%', '91%', 'Compliant'],
        ['Keyboard Navigation', '45%', '88%', 'Improving'],
        ['Screen Reader Support', '52%', '79%', 'Improving'],
        ['Form Accessibility', '71%', '94%', 'Compliant'],
    ]

    for row_idx, row_data in enumerate(data_rows, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell_run = cell.paragraphs[0].add_run(value)
            cell_run.font.name = 'Calibri'
            cell_run.font.size = Pt(10)

    # Footer note
    footer_note = doc.add_paragraph()
    footer_note.paragraph_format.space_before = Pt(16)
    footer_run = footer_note.add_run(
        'For questions regarding this report, contact the Accessibility Team at '
        'accessibility@acmecorp.com or call ext. 4892.'
    )
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(130, 130, 130)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Verify no alt text is present on the image
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    doc_check = Document(OUTPUT)
    alt_texts_found = []
    for para in doc_check.paragraphs:
        for run in para.runs:
            for docPr in run._element.iter(f'{{{wp_ns}}}docPr'):
                descr = docPr.get('descr', '')
                if descr:
                    alt_texts_found.append(descr)
    if alt_texts_found:
        print(f'WARNING: Alt text found in initial: {alt_texts_found}')
    else:
        print('VERIFIED: No alt text on image in initial file (correct)')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
