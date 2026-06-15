"""
Initial Setup: Delete the last stray test note line from a completed document
Task ID: writer_edit_067
Domain: libreoffice_writer

Creates /home/user/Desktop/final_draft.docx — a 5-page completed business partnership
document whose last paragraph is a stray editor test note that should be deleted.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_edit_067'
OUTPUT = f'{WORKDIR}/final_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Meridian Solutions Group  |  Confidential"
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = hp.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # --- Title block ---
    title = doc.add_heading("Strategic Partnership Report", level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph("Prepared for: Nexus Capital Partners")
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(12)
        run.italic = True

    date_para = doc.add_paragraph("Date: March 15, 2025")
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in date_para.runs:
        run.font.size = Pt(11)

    doc.add_paragraph("")  # spacer

    # --- Section 1: Executive Summary ---
    h1 = doc.add_heading("1. Executive Summary", level=1)

    p = doc.add_paragraph(
        "This report provides a comprehensive review of the strategic partnership between "
        "Meridian Solutions Group and Nexus Capital Partners over the fiscal year 2024–2025. "
        "Our collaboration has yielded significant results across joint marketing initiatives, "
        "technology co-development, and shared client acquisition programs."
    )
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(
        "Overall, the partnership delivered a combined revenue impact of $4.2 million, "
        "representing a 17% increase over the prior year. Client satisfaction scores improved "
        "from 81% to 89%, and the jointly developed analytics platform now serves over 340 "
        "enterprise accounts."
    )
    p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # --- Section 2: Key Milestones ---
    doc.add_heading("2. Key Milestones Achieved", level=1)

    milestones = [
        ("Q1 2025 — Platform Launch",
         "The Nexus-Meridian Analytics Hub was successfully launched in January 2025, "
         "onboarding 87 enterprise clients within the first 30 days. The platform exceeded "
         "initial traffic projections by 43%, requiring an unplanned but well-executed "
         "infrastructure scale-up."),
        ("Q2 2025 — Joint Marketing Campaign",
         "A co-branded digital marketing campaign reached 1.8 million impressions across "
         "LinkedIn, industry publications, and targeted email outreach. The campaign generated "
         "214 qualified leads, of which 62 converted to active accounts by quarter end."),
        ("Q3 2025 — Client Expansion",
         "Based on early platform success, both organizations agreed to expand the client "
         "coverage area to include mid-market segments. A dedicated support team of 12 "
         "professionals was established, reducing average issue resolution time from 48 hours "
         "to under 6 hours."),
        ("Q4 2025 — Renewal Negotiations",
         "Formal renewal discussions commenced in October. Both parties agreed in principle "
         "to a 3-year extension with expanded terms covering additional product verticals and "
         "geographic markets, subject to final legal review."),
    ]

    for title_text, body_text in milestones:
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(6)
        run_title = ph.add_run(title_text)
        run_title.bold = True
        run_title.font.size = Pt(11)

        pb = doc.add_paragraph(body_text)
        pb.paragraph_format.left_indent = Inches(0.25)
        pb.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # --- Section 3: Financial Review ---
    doc.add_heading("3. Financial Review", level=1)

    p = doc.add_paragraph(
        "The table below summarizes revenue contributions and cost allocations for the "
        "partnership activities in fiscal year 2024–2025."
    )
    p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers_data = ["Category", "Meridian (USD)", "Nexus (USD)", "Combined (USD)"]
    for col_idx, hdr in enumerate(headers_data):
        cell = table.cell(0, col_idx)
        cell.text = hdr
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    rows_data = [
        ["Licensing Revenue", "1,840,000", "1,200,000", "3,040,000"],
        ["Professional Services", "540,000", "310,000", "850,000"],
        ["Marketing Spend", "(220,000)", "(180,000)", "(400,000)"],
        ["Infrastructure Costs", "(95,000)", "(75,000)", "(170,000)"],
        ["Net Contribution", "2,065,000", "1,255,000", "3,320,000"],
    ]
    for row_idx, row_vals in enumerate(rows_data, 1):
        for col_idx, val in enumerate(row_vals):
            cell = table.cell(row_idx, col_idx)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                if row_idx == 5:
                    run.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph(
        "These figures represent direct partnership-attributable contributions only and exclude "
        "indirect benefits such as brand equity gains and talent development. A full audit "
        "report is available upon request from the Finance department of either organization."
    )
    p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # --- Section 4: Challenges and Resolutions ---
    doc.add_heading("4. Challenges and Resolutions", level=1)

    challenges = [
        ("Integration Complexity",
         "Early API integration between the two technology stacks encountered compatibility "
         "issues due to differing authentication frameworks. A dedicated integration team "
         "resolved these within 3 weeks, and the final solution was documented and shared "
         "as a reusable connector module for future platform extensions."),
        ("Data Privacy Compliance",
         "Expansion into European markets required additional GDPR compliance reviews. "
         "Both legal teams collaborated on a data processing agreement (DPA) that satisfied "
         "regulatory requirements while minimizing operational overhead. The approved DPA "
         "now serves as the template for all future EU-scope engagements."),
        ("Resource Allocation Disputes",
         "In Q2, both organizations identified overlapping resource assignments on the shared "
         "account management team. A revised RACI matrix was developed and implemented, "
         "clearly delineating responsibilities and eliminating duplication."),
    ]

    for title_text, body_text in challenges:
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(6)
        run_title = ph.add_run(title_text)
        run_title.bold = True
        run_title.font.size = Pt(11)

        pb = doc.add_paragraph(body_text)
        pb.paragraph_format.left_indent = Inches(0.25)
        pb.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # --- Section 5: Recommendations and Next Steps ---
    doc.add_heading("5. Recommendations and Next Steps", level=1)

    p = doc.add_paragraph(
        "Based on the outcomes of the 2024–2025 partnership cycle, Meridian Solutions Group "
        "recommends the following strategic priorities for the upcoming renewal period:"
    )
    p.paragraph_format.space_after = Pt(6)

    recommendations = [
        "Expand the Nexus-Meridian Analytics Hub to include predictive modeling modules by Q2 2026.",
        "Establish a joint innovation steering committee meeting quarterly to align product roadmaps.",
        "Increase co-marketing investment by 25% to capitalize on strong brand recognition in the enterprise segment.",
        "Develop a structured talent exchange program enabling 4–6 professionals annually to rotate between organizations.",
        "Formalize a shared IP ownership framework prior to commencement of any new co-development projects.",
    ]

    for rec in recommendations:
        p_item = doc.add_paragraph(rec, style="List Bullet")
        p_item.paragraph_format.space_after = Pt(4)

    doc.add_paragraph("")

    # --- Section 6: Closing ---
    doc.add_heading("6. Closing Remarks", level=1)

    p = doc.add_paragraph(
        "Meridian Solutions Group values the depth and quality of the relationship established "
        "with Nexus Capital Partners over the past year. The achievements documented in this "
        "report reflect the commitment, professionalism, and collaborative spirit demonstrated "
        "by both organizations at every level."
    )
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(
        "As we look ahead to an expanded partnership under the proposed renewal terms, we are "
        "confident that continued alignment on strategic goals will yield measurable and "
        "sustainable outcomes for both organizations and our shared client base."
    )
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(
        "We look forward to your continued partnership and support."
    )
    p.paragraph_format.space_after = Pt(12)

    # --- STRAY TEST NOTE (last line — must be deleted by agent) ---
    stray = doc.add_paragraph(
        "DELETE THIS - test note by editor, do not include in final version"
    )
    for run in stray.runs:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
