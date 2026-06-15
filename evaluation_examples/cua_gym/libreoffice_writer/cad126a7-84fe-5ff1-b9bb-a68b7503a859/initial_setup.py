"""
Initial Setup: Press release document with all paragraphs left-aligned
Task ID: writer_para_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'writer_para_058'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # All paragraphs are left-aligned (default) in the initial state.
    # The task asks the agent to apply center/right/justify alignment.

    paragraphs = [
        'FOR IMMEDIATE RELEASE',
        'GreenTech Corp Announces Breakthrough in Solar Panel Efficiency',
        'March 3, 2025',
        'SAN FRANCISCO \u2014 GreenTech Corporation today announced a major breakthrough in photovoltaic cell technology that increases solar panel efficiency by 35% compared to current commercial panels.',
        'The new technology, developed at the company research facility in collaboration with MIT, uses a novel perovskite-silicon tandem structure that captures a broader spectrum of sunlight.',
        'CEO Dr. Amanda Liu stated that commercial production is expected to begin in Q3 2025, with initial deployments planned for utility-scale solar farms in California and Arizona.',
        'About GreenTech Corporation',
        'GreenTech Corporation is a leading innovator in renewable energy technology, headquartered in San Francisco with operations in 12 countries.',
    ]

    for text in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        # All paragraphs explicitly set to LEFT alignment in initial state
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
