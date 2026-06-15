"""
Reward Script: Create a pricing table in Proposal_v2.docx
Task ID: writer_pd_015
Domain: libreoffice_writer
Scoring:
  Component 1: Pricing table existence & dimensions (0.20)
  Component 2: Header row content (0.15)
  Component 3: Header row formatting - dark blue bg, white bold text (0.15)
  Component 4: 8 service line items present (0.15)
  Component 5: Summary rows present with correct labels (0.15)
  Component 6: Numeric columns right-aligned (0.10)
  Component 7: Mathematical consistency of values (0.10)
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_015'

def parse_currency(text):
    """Parse currency string like '$59,200.00' or '-$62,160.00' to float."""
    if not text:
        return None
    try:
        cleaned = text.strip().replace('$', '').replace(',', '')
        return float(cleaned)
    except (ValueError, AttributeError):
        return None

def get_cell_shading(cell):
    """Get background fill color from cell shading element."""
    tc_pr = cell._element.find(qn('w:tcPr'))
    if tc_pr is None:
        return None
    shading = tc_pr.find(qn('w:shd'))
    if shading is None:
        return None
    return shading.get(qn('w:fill'))

def has_bold_text(cell):
    """Check if cell contains bold text (via run property or XML element)."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for para in cell.paragraphs:
        for rpr in para._element.findall('.//w:rPr', ns):
            b = rpr.find(qn('w:b'))
            if b is not None:
                # w:b present without val or with val != "0" means bold
                val = b.get(qn('w:val'))
                if val is None or val not in ('0', 'false'):
                    return True
        for run in para.runs:
            if run.font.bold:
                return True
    return False

def has_white_text(cell):
    """Check if cell text color is white (FFFFFF)."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for para in cell.paragraphs:
        for rpr in para._element.findall('.//w:rPr', ns):
            color = rpr.find(qn('w:color'))
            if color is not None:
                val = color.get(qn('w:val'))
                if val and val.upper() == 'FFFFFF':
                    return True
        for run in para.runs:
            if run.font.color and run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF):
                return True
    return False

def is_right_aligned(cell):
    """Check if the first paragraph in the cell is right-aligned."""
    for para in cell.paragraphs:
        if para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return True
    return False

def find_pricing_table(doc):
    """Find the pricing table - it should have 5 columns with header: Service, Unit, Quantity, Unit Price, Total.
    Returns the table object or None."""
    expected_headers = ['service', 'unit', 'quantity', 'unit price', 'total']
    for table in doc.tables:
        if len(table.columns) == 5 and len(table.rows) >= 5:
            header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if header_texts == expected_headers:
                return table
    # Fallback: any table with 5 columns that isn't the timeline table
    for table in doc.tables:
        if len(table.columns) == 5:
            first_cell = table.rows[0].cells[0].text.strip().lower()
            if first_cell != 'phase':
                return table
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the pricing table (must be a NEW table, not the pre-existing timeline table)
    pricing_table = find_pricing_table(doc)

    # Component 1: Pricing table existence & dimensions (0.20 points)
    try:
        if pricing_table is None:
            print("FAIL: Component 1 — No pricing table found with 5 columns and correct headers")
            print("REWARD: 0.0")
            return 0.0

        num_rows = len(pricing_table.rows)
        num_cols = len(pricing_table.columns)

        if num_cols == 5 and num_rows >= 12:
            print(f"PASS: Component 1 — Pricing table found: {num_rows} rows x {num_cols} cols (0.20 pts)")
            total_score += 0.20
        elif num_cols == 5 and num_rows >= 10:
            print(f"PARTIAL: Component 1 — Table has {num_rows} rows (expected >= 12), {num_cols} cols (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Table dimensions: {num_rows} rows x {num_cols} cols, expected >= 12 rows x 5 cols")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row content (0.15 points)
    try:
        header_texts = [cell.text.strip() for cell in pricing_table.rows[0].cells]
        expected = ['Service', 'Unit', 'Quantity', 'Unit Price', 'Total']
        matches = sum(1 for a, b in zip(header_texts, expected) if a.lower() == b.lower())
        if matches == 5:
            print(f"PASS: Component 2 — Header row has all 5 correct column names: {header_texts} (0.15 pts)")
            total_score += 0.15
        elif matches >= 3:
            partial = 0.15 * (matches / 5)
            print(f"PARTIAL: Component 2 — {matches}/5 headers match: {header_texts} ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Expected headers {expected}, found {header_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row formatting - dark blue background, white bold text (0.15 points)
    try:
        header_cells = pricing_table.rows[0].cells
        blue_count = 0
        bold_count = 0
        white_count = 0
        for cell in header_cells:
            fill = get_cell_shading(cell)
            if fill and fill.upper() in ('003366', '003399', '002244', '003355', '002255'):
                blue_count += 1
            elif fill and fill.upper().startswith('00') and len(fill) == 6:
                # Accept any dark blue-ish shade
                r_val = int(fill[0:2], 16)
                g_val = int(fill[2:4], 16)
                b_val = int(fill[4:6], 16)
                if r_val < 20 and g_val < 80 and b_val > 80:
                    blue_count += 1
            if has_bold_text(cell):
                bold_count += 1
            if has_white_text(cell):
                white_count += 1

        sub_score = 0.0
        if blue_count >= 4:
            sub_score += 0.05
            print(f"  PASS: Header shading dark blue ({blue_count}/5 cells)")
        else:
            print(f"  FAIL: Header shading — expected dark blue, found {blue_count}/5 cells with correct fill")
        if bold_count >= 4:
            sub_score += 0.05
            print(f"  PASS: Header bold text ({bold_count}/5 cells)")
        else:
            print(f"  FAIL: Header bold — {bold_count}/5 cells")
        if white_count >= 4:
            sub_score += 0.05
            print(f"  PASS: Header white text ({white_count}/5 cells)")
        else:
            print(f"  FAIL: Header white text — {white_count}/5 cells")

        if sub_score > 0:
            print(f"PASS: Component 3 — Header formatting ({sub_score:.2f} pts)")
            total_score += sub_score
        else:
            print(f"FAIL: Component 3 — No header formatting detected")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 8 service line items present (0.15 points)
    try:
        # Check rows 1-8 have service content (non-empty first cell and last cell with $)
        service_count = 0
        for r_idx in range(1, min(9, len(pricing_table.rows))):
            row = pricing_table.rows[r_idx]
            service_name = row.cells[0].text.strip()
            total_val = row.cells[4].text.strip() if len(row.cells) > 4 else ''
            if service_name and '$' in total_val:
                service_count += 1

        if service_count == 8:
            print(f"PASS: Component 4 — All 8 service line items present (0.15 pts)")
            total_score += 0.15
        elif service_count >= 5:
            partial = 0.15 * (service_count / 8)
            print(f"PARTIAL: Component 4 — {service_count}/8 service line items ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {service_count}/8 service line items found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Summary rows with correct labels (0.15 points)
    try:
        # Look for Subtotal, Discount, Tax, Grand Total in the table
        all_first_cells = [row.cells[0].text.strip().lower() for row in pricing_table.rows]

        found_subtotal = any('subtotal' in c for c in all_first_cells)
        found_discount = any('discount' in c for c in all_first_cells)
        found_tax = any('tax' in c for c in all_first_cells)
        found_grand = any('grand total' in c for c in all_first_cells)

        summary_count = sum([found_subtotal, found_discount, found_tax, found_grand])

        if summary_count == 4:
            print(f"PASS: Component 5 — All 4 summary rows present: Subtotal, Discount, Tax, Grand Total (0.15 pts)")
            total_score += 0.15
        elif summary_count >= 2:
            partial = 0.15 * (summary_count / 4)
            print(f"PARTIAL: Component 5 — {summary_count}/4 summary rows found ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {summary_count}/4 summary rows found")
            print(f"  Labels found: {all_first_cells}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Numeric columns right-aligned (0.10 points)
    try:
        # Check columns 2 (Quantity), 3 (Unit Price), 4 (Total) for right-alignment
        # Check across data rows (1-8) and header row
        right_aligned_count = 0
        total_checks = 0
        for r_idx in range(0, min(9, len(pricing_table.rows))):
            for c_idx in [2, 3, 4]:
                if c_idx < len(pricing_table.rows[r_idx].cells):
                    cell = pricing_table.rows[r_idx].cells[c_idx]
                    if cell.text.strip():  # only check non-empty cells
                        total_checks += 1
                        if is_right_aligned(cell):
                            right_aligned_count += 1

        if total_checks > 0:
            ratio = right_aligned_count / total_checks
            if ratio >= 0.8:
                print(f"PASS: Component 6 — {right_aligned_count}/{total_checks} numeric cells right-aligned (0.10 pts)")
                total_score += 0.10
            elif ratio >= 0.5:
                partial = 0.10 * ratio
                print(f"PARTIAL: Component 6 — {right_aligned_count}/{total_checks} numeric cells right-aligned ({partial:.2f} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 6 — Only {right_aligned_count}/{total_checks} numeric cells right-aligned")
        else:
            print(f"FAIL: Component 6 — No numeric cells found to check alignment")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Mathematical consistency (0.10 points)
    try:
        # Extract values from the table and verify math
        item_totals = []
        for r_idx in range(1, min(9, len(pricing_table.rows))):
            row = pricing_table.rows[r_idx]
            qty_text = row.cells[2].text.strip() if len(row.cells) > 2 else ''
            price_text = row.cells[3].text.strip() if len(row.cells) > 3 else ''
            total_text = row.cells[4].text.strip() if len(row.cells) > 4 else ''

            qty = parse_currency(qty_text)
            price = parse_currency(price_text)
            total = parse_currency(total_text)

            if qty is not None and price is not None and total is not None:
                item_totals.append(total)
                expected_total = qty * price
                if abs(total - expected_total) > 0.02:
                    print(f"  WARN: Row {r_idx} math: {qty} x {price} = {expected_total}, found {total}")

        # Find summary row values
        subtotal_val = None
        discount_val = None
        tax_val = None
        grand_total_val = None

        for row in pricing_table.rows:
            label = row.cells[0].text.strip().lower()
            last_cell_text = row.cells[-1].text.strip()
            val = parse_currency(last_cell_text)

            if 'subtotal' in label and 'grand' not in label:
                subtotal_val = val
            elif 'discount' in label:
                discount_val = val
            elif 'tax' in label:
                tax_val = val
            elif 'grand total' in label:
                grand_total_val = val

        math_checks_passed = 0
        math_checks_total = 0

        # Check subtotal = sum of item totals
        if subtotal_val is not None and len(item_totals) > 0:
            math_checks_total += 1
            expected_subtotal = sum(item_totals)
            if abs(subtotal_val - expected_subtotal) < 1.0:
                math_checks_passed += 1
                print(f"  PASS: Subtotal {subtotal_val} matches sum of items {expected_subtotal}")
            else:
                print(f"  FAIL: Subtotal {subtotal_val} != sum of items {expected_subtotal}")

        # Check discount = -10% of subtotal
        if discount_val is not None and subtotal_val is not None:
            math_checks_total += 1
            expected_discount = -subtotal_val * 0.10
            if abs(discount_val - expected_discount) < 1.0:
                math_checks_passed += 1
                print(f"  PASS: Discount {discount_val} matches -10% of subtotal {expected_discount}")
            else:
                print(f"  FAIL: Discount {discount_val} != -10% of subtotal {expected_discount}")

        # Check grand total consistency
        if grand_total_val is not None and subtotal_val is not None and discount_val is not None and tax_val is not None:
            math_checks_total += 1
            expected_grand = subtotal_val + discount_val + tax_val
            if abs(grand_total_val - expected_grand) < 1.0:
                math_checks_passed += 1
                print(f"  PASS: Grand total {grand_total_val} matches {expected_grand}")
            else:
                print(f"  FAIL: Grand total {grand_total_val} != {expected_grand}")

        if math_checks_total > 0:
            ratio = math_checks_passed / math_checks_total
            score = 0.10 * ratio
            if ratio >= 0.9:
                print(f"PASS: Component 7 — Math consistency: {math_checks_passed}/{math_checks_total} checks passed (0.10 pts)")
                total_score += 0.10
            elif ratio > 0:
                print(f"PARTIAL: Component 7 — Math consistency: {math_checks_passed}/{math_checks_total} checks ({score:.2f} pts)")
                total_score += score
            else:
                print(f"FAIL: Component 7 — Math consistency: 0/{math_checks_total} checks passed")
        else:
            print(f"FAIL: Component 7 — Could not extract enough values to verify math")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")

# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
