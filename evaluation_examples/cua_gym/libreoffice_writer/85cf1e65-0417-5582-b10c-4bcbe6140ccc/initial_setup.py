"""
Initial Setup: Create a 10-page proposal document with 'Pricing Schedule' heading on page 8.
Task ID: writer_pd_015
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add an explicit page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def set_heading_style(para, level=1):
    """Style a heading paragraph."""
    for run in para.runs:
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)


def create_initial():
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== PAGE 1: Title Page =====
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading('Proposal_v2', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Digital Transformation Initiative')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info.add_run('Prepared for: Meridian Healthcare Group\nPrepared by: Nexus Consulting Partners\nDate: March 15, 2025\nVersion: 2.0')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_page_break(doc)

    # ===== PAGE 2: Table of Contents =====
    h = doc.add_heading('Table of Contents', level=1)
    set_heading_style(h)
    doc.add_paragraph()

    toc_items = [
        ('1.', 'Executive Summary', '3'),
        ('2.', 'Company Background', '4'),
        ('3.', 'Project Scope & Objectives', '5'),
        ('4.', 'Methodology & Approach', '6'),
        ('5.', 'Project Timeline', '7'),
        ('6.', 'Pricing Schedule', '8'),
        ('7.', 'Terms & Conditions', '9'),
        ('8.', 'Appendices', '10'),
    ]
    for num, title_text, page in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(f'{num}  {title_text}')
        r.font.size = Pt(12)
        r2 = p.add_run(f'  {"." * 60}  {page}')
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    add_page_break(doc)

    # ===== PAGE 3: Executive Summary =====
    h = doc.add_heading('1. Executive Summary', level=1)
    set_heading_style(h)
    doc.add_paragraph(
        'Nexus Consulting Partners is pleased to present this revised proposal for the '
        'Digital Transformation Initiative at Meridian Healthcare Group. This comprehensive '
        'engagement will modernize your patient management systems, streamline clinical workflows, '
        'and enhance data analytics capabilities across all 12 regional facilities.'
    )
    doc.add_paragraph(
        'Our team of 25 certified consultants brings extensive experience in healthcare IT '
        'transformation, having successfully delivered similar projects for Providence Health, '
        'Kaiser Permanente, and Cleveland Clinic. We propose a phased approach spanning 18 months, '
        'designed to minimize operational disruption while maximizing ROI.'
    )
    doc.add_paragraph(
        'Key deliverables include: a unified Electronic Health Records (EHR) platform, real-time '
        'analytics dashboards for clinical and operational metrics, automated compliance reporting '
        'for HIPAA and state regulations, and a comprehensive training program for over 3,000 staff members.'
    )
    doc.add_paragraph(
        'The estimated total investment for this initiative is $2,847,500, with projected annual '
        'savings of $1.2 million in operational efficiencies and reduced compliance costs beginning '
        'in Year 2 of implementation.'
    )

    add_page_break(doc)

    # ===== PAGE 4: Company Background =====
    h = doc.add_heading('2. Company Background', level=1)
    set_heading_style(h)
    doc.add_heading('2.1 About Nexus Consulting Partners', level=2)
    doc.add_paragraph(
        'Founded in 2008, Nexus Consulting Partners has grown from a boutique technology advisory '
        'firm to a nationally recognized leader in healthcare digital transformation. Headquartered '
        'in Chicago, Illinois, we maintain regional offices in Boston, San Francisco, and Atlanta.'
    )
    doc.add_paragraph(
        'Our healthcare practice employs over 150 specialists, including certified Epic consultants, '
        'Cerner implementation experts, cloud architects, and cybersecurity professionals. We hold '
        'partnerships with Microsoft, AWS, Salesforce Health Cloud, and leading EHR vendors.'
    )
    doc.add_heading('2.2 Relevant Experience', level=2)
    doc.add_paragraph(
        'Providence Health Systems (2022-2023): Led a $4.5M EHR migration affecting 18 hospitals '
        'and 45,000 clinical staff. Delivered on time and 8% under budget.'
    )
    doc.add_paragraph(
        'Bay Area Medical Center (2023-2024): Implemented real-time patient flow analytics, '
        'reducing average ER wait times by 34% and improving bed utilization by 22%.'
    )
    doc.add_paragraph(
        'Heartland Regional Health Network (2024): Deployed automated HIPAA compliance monitoring '
        'across 6 facilities, reducing audit preparation time by 60%.'
    )

    add_page_break(doc)

    # ===== PAGE 5: Project Scope & Objectives =====
    h = doc.add_heading('3. Project Scope & Objectives', level=1)
    set_heading_style(h)
    doc.add_heading('3.1 Scope of Work', level=2)
    doc.add_paragraph(
        'This engagement encompasses the full lifecycle of digital transformation for Meridian '
        'Healthcare Group, including assessment, planning, implementation, testing, deployment, '
        'and post-launch support. The scope covers all 12 regional facilities and the central '
        'administrative office.'
    )

    doc.add_heading('3.2 Primary Objectives', level=2)
    objectives = [
        'Consolidate disparate clinical systems into a unified EHR platform',
        'Implement real-time operational and clinical analytics dashboards',
        'Automate regulatory compliance reporting (HIPAA, state-level requirements)',
        'Develop and execute a comprehensive change management program',
        'Establish a secure, cloud-based infrastructure with 99.99% uptime SLA',
        'Create standardized clinical workflows across all facilities',
        'Deploy patient engagement portal with telehealth capabilities',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('3.3 Out of Scope', level=2)
    doc.add_paragraph(
        'The following items are explicitly excluded from this engagement: hardware procurement '
        'and installation, third-party software licensing fees, physical network infrastructure '
        'upgrades, and ongoing managed services beyond the 90-day post-launch support period.'
    )

    add_page_break(doc)

    # ===== PAGE 6: Methodology & Approach =====
    h = doc.add_heading('4. Methodology & Approach', level=1)
    set_heading_style(h)
    doc.add_paragraph(
        'Our proven TRANSFORM methodology ensures successful delivery through structured phases, '
        'continuous stakeholder engagement, and rigorous quality gates at each milestone.'
    )

    doc.add_heading('Phase 1: Discovery & Assessment (Weeks 1-8)', level=2)
    doc.add_paragraph(
        'Comprehensive audit of existing systems, workflows, and data architecture. Stakeholder '
        'interviews with department heads, clinical staff, and IT leadership. Gap analysis and '
        'requirements documentation. Deliverables: Current State Assessment Report, Requirements '
        'Specification Document, Risk Register.'
    )
    doc.add_heading('Phase 2: Design & Architecture (Weeks 9-16)', level=2)
    doc.add_paragraph(
        'Solution architecture design, technology stack selection, and integration mapping. '
        'Prototype development for key user interfaces. Security architecture review and '
        'compliance framework alignment. Deliverables: Solution Design Document, Architecture '
        'Diagrams, Security Plan.'
    )
    doc.add_heading('Phase 3: Build & Configure (Weeks 17-40)', level=2)
    doc.add_paragraph(
        'Iterative development sprints with bi-weekly stakeholder demos. EHR configuration and '
        'customization, data migration planning and execution, analytics dashboard development. '
        'Integration with existing systems (billing, scheduling, lab). Deliverables: Configured '
        'EHR System, Analytics Platform, Integration Documentation.'
    )
    doc.add_heading('Phase 4: Testing & Validation (Weeks 41-52)', level=2)
    doc.add_paragraph(
        'Comprehensive UAT with clinical and administrative users. Performance testing under '
        'simulated peak loads. Security penetration testing and compliance validation. '
        'Deliverables: Test Results Report, Compliance Certification, Go-Live Readiness Assessment.'
    )

    add_page_break(doc)

    # ===== PAGE 7: Project Timeline =====
    h = doc.add_heading('5. Project Timeline', level=1)
    set_heading_style(h)
    doc.add_paragraph(
        'The following timeline outlines the major milestones and delivery dates for the '
        'Digital Transformation Initiative. All dates assume a project start date of Q2 2025.'
    )

    # Timeline table
    timeline = doc.add_table(rows=9, cols=4)
    timeline.style = 'Table Grid'
    headers = ['Phase', 'Duration', 'Start Date', 'End Date']
    for i, hdr in enumerate(headers):
        cell = timeline.cell(0, i)
        cell.text = hdr
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows_data = [
        ['Discovery & Assessment', '8 weeks', 'Apr 2025', 'May 2025'],
        ['Design & Architecture', '8 weeks', 'Jun 2025', 'Jul 2025'],
        ['Build & Configure - Phase A', '12 weeks', 'Aug 2025', 'Oct 2025'],
        ['Build & Configure - Phase B', '12 weeks', 'Nov 2025', 'Jan 2026'],
        ['Testing & Validation', '12 weeks', 'Feb 2026', 'Apr 2026'],
        ['Pilot Deployment (3 facilities)', '4 weeks', 'May 2026', 'May 2026'],
        ['Full Rollout', '8 weeks', 'Jun 2026', 'Jul 2026'],
        ['Post-Launch Support', '12 weeks', 'Aug 2026', 'Oct 2026'],
    ]
    for r, row_data in enumerate(rows_data, 1):
        for c, val in enumerate(row_data):
            timeline.cell(r, c).text = val

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Timeline assumes timely completion of prerequisites including vendor contract '
        'finalization, infrastructure provisioning, and staff availability for training sessions.'
    )

    add_page_break(doc)

    # ===== PAGE 8: Pricing Schedule (NO TABLE - task asks agent to create it) =====
    h = doc.add_heading('6. Pricing Schedule', level=1)
    set_heading_style(h)
    doc.add_paragraph(
        'The following pricing reflects our competitive rates for the scope of services described '
        'in this proposal. All fees are quoted in US dollars and are valid for 90 days from the '
        'date of this proposal.'
    )
    # NOTE: No pricing table here. The task asks the agent to create it.

    add_page_break(doc)

    # ===== PAGE 9: Terms & Conditions =====
    h = doc.add_heading('7. Terms & Conditions', level=1)
    set_heading_style(h)

    doc.add_heading('7.1 Payment Terms', level=2)
    doc.add_paragraph(
        'Payment shall be made in four installments: 25% upon contract execution, 25% at the '
        'completion of Phase 2 (Design & Architecture), 25% at the completion of Phase 3 '
        '(Build & Configure), and 25% upon final acceptance and go-live approval.'
    )

    doc.add_heading('7.2 Intellectual Property', level=2)
    doc.add_paragraph(
        'All custom-developed software, configurations, and documentation created specifically '
        'for Meridian Healthcare Group during this engagement shall become the exclusive property '
        'of the Client upon final payment. Nexus Consulting Partners retains rights to its '
        'pre-existing frameworks, tools, and methodologies.'
    )

    doc.add_heading('7.3 Confidentiality', level=2)
    doc.add_paragraph(
        'Both parties agree to maintain strict confidentiality of all proprietary information, '
        'patient data, and trade secrets disclosed during this engagement. This obligation '
        'survives termination of the agreement for a period of five (5) years.'
    )

    doc.add_heading('7.4 Limitation of Liability', level=2)
    doc.add_paragraph(
        'The total aggregate liability of Nexus Consulting Partners under this agreement shall '
        'not exceed the total fees paid by the Client. Neither party shall be liable for '
        'indirect, incidental, or consequential damages.'
    )

    add_page_break(doc)

    # ===== PAGE 10: Appendices =====
    h = doc.add_heading('8. Appendices', level=1)
    set_heading_style(h)

    doc.add_heading('Appendix A: Team Biographies', level=2)
    doc.add_paragraph(
        'Dr. Elena Vasquez, Project Director - 18 years of healthcare IT experience. Former CIO '
        'of Regional Medical Center. Certified PMP, ITIL v4, and Epic Certified Consultant.'
    )
    doc.add_paragraph(
        'James Thornton, Lead Architect - 15 years in enterprise architecture. AWS Solutions '
        'Architect Professional. Led cloud migrations for 20+ healthcare organizations.'
    )
    doc.add_paragraph(
        'Dr. Priya Sharma, Clinical Informatics Lead - Board-certified in clinical informatics. '
        'Former CMIO at University Hospital. Published researcher in health information exchange.'
    )

    doc.add_heading('Appendix B: Client References', level=2)
    doc.add_paragraph(
        'Providence Health Systems - Contact: Sarah Mitchell, VP of Information Services. '
        'Phone: (503) 555-0147. Email: s.mitchell@providence.org'
    )
    doc.add_paragraph(
        'Bay Area Medical Center - Contact: Dr. Robert Kim, Chief Medical Officer. '
        'Phone: (415) 555-0283. Email: r.kim@bayareamc.org'
    )

    doc.add_heading('Appendix C: Certifications & Compliance', level=2)
    doc.add_paragraph(
        'Nexus Consulting Partners maintains the following certifications relevant to this '
        'engagement: SOC 2 Type II, ISO 27001, HITRUST CSF, and FedRAMP Moderate authorization. '
        'All consultants undergo annual HIPAA training and background checks.'
    )

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
