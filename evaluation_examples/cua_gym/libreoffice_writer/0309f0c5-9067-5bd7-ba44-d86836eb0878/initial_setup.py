"""
Initial Setup: Create a Software Spec document with 15 tracked changes from 3 authors.
Task ID: writer_rm_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import datetime

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Tracked change authors
AUTHORS = [
    "Sarah Chen",
    "Marcus Rivera",
    "Elena Kowalski",
]

# Base dates for changes (over the past week)
BASE_DATES = [
    "2026-03-26T09:15:00Z",
    "2026-03-26T14:30:00Z",
    "2026-03-27T10:45:00Z",
    "2026-03-27T11:20:00Z",
    "2026-03-27T16:00:00Z",
    "2026-03-28T08:30:00Z",
    "2026-03-28T13:15:00Z",
    "2026-03-29T09:00:00Z",
    "2026-03-29T10:30:00Z",
    "2026-03-29T15:45:00Z",
    "2026-03-30T11:00:00Z",
    "2026-03-30T14:20:00Z",
    "2026-03-31T09:30:00Z",
    "2026-03-31T16:45:00Z",
    "2026-04-01T10:00:00Z",
]

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}


def make_run_element(text, bold=False, font_name="Calibri", font_size_pt=11):
    """Create a w:r element with optional formatting."""
    w = NSMAP['w']
    r_elem = etree.SubElement(etree.Element('dummy'), qn('w:r'))
    rPr = etree.SubElement(r_elem, qn('w:rPr'))
    if bold:
        etree.SubElement(rPr, qn('w:b'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(font_size_pt * 2))  # half-points
    t_elem = etree.SubElement(r_elem, qn('w:t'))
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_elem.text = text
    return r_elem


def make_del_run_element(text, font_name="Calibri", font_size_pt=11):
    """Create a w:r element with w:delText for deletions."""
    w = NSMAP['w']
    r_elem = etree.SubElement(etree.Element('dummy'), qn('w:r'))
    rPr = etree.SubElement(r_elem, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(font_size_pt * 2))
    t_elem = etree.SubElement(r_elem, qn('w:delText'))
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_elem.text = text
    return r_elem


def add_insertion(para_elem, text, author, date, rev_id):
    """Add a tracked insertion (w:ins) to a paragraph element."""
    ins = etree.SubElement(para_elem, qn('w:ins'))
    ins.set(qn('w:id'), str(rev_id))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)
    run = make_run_element(text)
    ins.append(run)
    return ins


def add_deletion(para_elem, text, author, date, rev_id):
    """Add a tracked deletion (w:del) to a paragraph element."""
    del_elem = etree.SubElement(para_elem, qn('w:del'))
    del_elem.set(qn('w:id'), str(rev_id))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date)
    run = make_del_run_element(text)
    del_elem.append(run)
    return del_elem


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # === Title ===
    title = doc.add_heading('Software Requirements Specification', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Project Atlas - Cloud Migration Platform')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()  # spacer

    # Version info
    ver = doc.add_paragraph()
    ver.add_run('Version: ').bold = True
    ver.add_run('2.3.1')
    date_p = doc.add_paragraph()
    date_p.add_run('Last Updated: ').bold = True
    date_p.add_run('April 1, 2026')
    authors_p = doc.add_paragraph()
    authors_p.add_run('Authors: ').bold = True
    authors_p.add_run('Sarah Chen, Marcus Rivera, Elena Kowalski')

    doc.add_paragraph()  # spacer

    # === Section 1: Introduction ===
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This document outlines the functional and non-functional requirements for '
        'Project Atlas, a comprehensive cloud migration platform designed to facilitate '
        'enterprise-scale infrastructure transitions from on-premises data centers to '
        'multi-cloud environments.'
    )
    doc.add_paragraph(
        'The platform supports automated workload assessment, dependency mapping, '
        'migration planning, and execution monitoring across AWS, Azure, and Google Cloud.'
    )

    # === Section 2: Scope ===
    doc.add_heading('2. Scope', level=1)
    doc.add_paragraph(
        'Project Atlas encompasses the following core modules:'
    )
    doc.add_paragraph('Discovery Engine - Automated infrastructure scanning and inventory', style='List Bullet')
    doc.add_paragraph('Dependency Mapper - Application dependency analysis and visualization', style='List Bullet')
    doc.add_paragraph('Migration Planner - Wave planning and resource allocation', style='List Bullet')
    doc.add_paragraph('Execution Controller - Orchestrated migration with rollback capability', style='List Bullet')
    doc.add_paragraph('Monitoring Dashboard - Real-time progress tracking and alerting', style='List Bullet')

    # === Section 3: Functional Requirements ===
    doc.add_heading('3. Functional Requirements', level=1)

    doc.add_heading('3.1 Discovery Engine', level=2)
    doc.add_paragraph(
        'The Discovery Engine shall perform automated scanning of on-premises infrastructure '
        'using agentless discovery protocols. It must support VMware vSphere, Microsoft Hyper-V, '
        'and bare-metal server environments.'
    )
    doc.add_paragraph(
        'Scan results shall include CPU utilization metrics, memory consumption patterns, '
        'storage I/O profiles, and network traffic analysis collected over a minimum 30-day period.'
    )

    doc.add_heading('3.2 Dependency Mapper', level=2)
    doc.add_paragraph(
        'The Dependency Mapper shall automatically identify application-level dependencies '
        'through network flow analysis and process-level monitoring. Dependencies shall be '
        'classified as critical, important, or optional.'
    )
    doc.add_paragraph(
        'The module must generate interactive dependency graphs exportable in SVG and PDF formats. '
        'Graph layouts shall support hierarchical, force-directed, and circular arrangements.'
    )

    doc.add_heading('3.3 Migration Planner', level=2)
    doc.add_paragraph(
        'The Migration Planner shall support wave-based migration scheduling with configurable '
        'parallelism constraints. Each wave must respect dependency ordering and resource capacity '
        'limits defined by the target cloud environment.'
    )
    doc.add_paragraph(
        'Cost estimation shall be provided for each migration wave, including compute, storage, '
        'network egress, and licensing costs across all supported cloud providers.'
    )

    doc.add_heading('3.4 Execution Controller', level=2)
    doc.add_paragraph(
        'The Execution Controller shall orchestrate migration tasks with configurable retry '
        'policies and automatic rollback on failure. Maximum rollback time shall not exceed '
        '15 minutes for any single workload.'
    )

    # === Section 4: Non-Functional Requirements ===
    doc.add_heading('4. Non-Functional Requirements', level=1)

    doc.add_heading('4.1 Performance', level=2)
    doc.add_paragraph(
        'The platform shall support concurrent migration of up to 500 workloads without '
        'degradation in monitoring responsiveness. Dashboard refresh latency must remain '
        'below 3 seconds under peak load conditions.'
    )

    doc.add_heading('4.2 Security', level=2)
    doc.add_paragraph(
        'All data in transit shall be encrypted using TLS 1.3. Credentials for cloud provider '
        'APIs shall be stored in HashiCorp Vault with automatic rotation every 90 days. '
        'Role-based access control shall enforce least-privilege principles.'
    )

    doc.add_heading('4.3 Availability', level=2)
    doc.add_paragraph(
        'The platform shall maintain 99.9% uptime for the monitoring and control plane. '
        'Disaster recovery shall support cross-region failover with RPO of 5 minutes and '
        'RTO of 30 minutes.'
    )

    # === Section 5: Technical Architecture ===
    doc.add_heading('5. Technical Architecture', level=1)
    doc.add_paragraph(
        'The platform is built on a microservices architecture deployed on Kubernetes. '
        'Core services communicate via gRPC with REST API gateways for external integration. '
        'Event-driven workflows use Apache Kafka for message brokering.'
    )

    # Add a table for component overview
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Component', 'Technology', 'Scaling', 'Owner']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    table_data = [
        ['Discovery Service', 'Python / FastAPI', 'Horizontal (3-10 pods)', 'Sarah Chen'],
        ['Dependency Analyzer', 'Go / gRPC', 'Horizontal (2-8 pods)', 'Marcus Rivera'],
        ['Migration Orchestrator', 'Java / Spring Boot', 'Vertical + Horizontal', 'Elena Kowalski'],
        ['Dashboard Frontend', 'React / TypeScript', 'CDN + 2 pods', 'Sarah Chen'],
        ['Data Pipeline', 'Apache Spark', 'Auto-scaling cluster', 'Marcus Rivera'],
    ]
    for r, row_data in enumerate(table_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()  # spacer

    # === Section 6: Timeline ===
    doc.add_heading('6. Project Timeline', level=1)
    doc.add_paragraph(
        'Phase 1 (Q1 2026): Discovery Engine and Dependency Mapper - Complete')
    doc.add_paragraph(
        'Phase 2 (Q2 2026): Migration Planner and basic Execution Controller')
    doc.add_paragraph(
        'Phase 3 (Q3 2026): Advanced orchestration, monitoring dashboard, and GA release')

    # Save the base document first
    doc.save(OUTPUT)

    # Now manipulate the XML to add tracked changes
    # Re-open and work with the XML directly
    from docx import Document as DocxDocument
    doc = DocxDocument(OUTPUT)
    body = doc.element.body

    # We need to add tracked changes across the document.
    # We'll modify specific paragraphs to include insertions and deletions.

    paragraphs = body.findall(qn('w:p'))
    rev_id = 100  # starting revision ID

    # Change 1: Insertion in intro paragraph - Sarah Chen added "comprehensive"
    # (already in text, we'll restructure as tracked insertion)
    # Instead, we'll add tracked changes to existing paragraphs by modifying the XML

    # Strategy: Add new tracked change elements to specific paragraphs
    # We'll append insertions and deletions to various paragraphs

    # Find paragraphs by index (skip heading paragraphs)
    all_paras = list(paragraphs)

    # Change 1: Sarah Chen - Insert text after intro paragraph
    # Add an insertion to paragraph index ~6 (first intro para)
    if len(all_paras) > 6:
        p = all_paras[6]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[0])
        ins.set(qn('w:date'), BASE_DATES[0])
        r = make_run_element(' including hybrid cloud configurations')
        ins.append(r)

    # Change 2: Marcus Rivera - Delete text in scope section
    if len(all_paras) > 8:
        p = all_paras[8]
        del_elem = etree.SubElement(p, qn('w:del'))
        del_elem.set(qn('w:id'), str(rev_id)); rev_id += 1
        del_elem.set(qn('w:author'), AUTHORS[1])
        del_elem.set(qn('w:date'), BASE_DATES[1])
        r = make_del_run_element(' and legacy mainframe systems')
        del_elem.append(r)

    # Change 3: Elena Kowalski - Insert requirement text
    if len(all_paras) > 14:
        p = all_paras[14]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[2])
        ins.set(qn('w:date'), BASE_DATES[2])
        r = make_run_element(' Custom filtering shall allow users to focus on specific dependency chains.')
        ins.append(r)

    # Change 4: Sarah Chen - Delete old performance number
    if len(all_paras) > 19:
        p = all_paras[19]
        del_elem = etree.SubElement(p, qn('w:del'))
        del_elem.set(qn('w:id'), str(rev_id)); rev_id += 1
        del_elem.set(qn('w:author'), AUTHORS[0])
        del_elem.set(qn('w:date'), BASE_DATES[3])
        r = make_del_run_element(' (previously 200 workloads)')
        del_elem.append(r)

    # Change 5: Marcus Rivera - Insert security clarification
    if len(all_paras) > 20:
        p = all_paras[20]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[1])
        ins.set(qn('w:date'), BASE_DATES[4])
        r = make_run_element(' Multi-factor authentication is mandatory for all administrative operations.')
        ins.append(r)

    # Change 6: Elena Kowalski - Delete outdated availability info
    if len(all_paras) > 21:
        p = all_paras[21]
        del_elem = etree.SubElement(p, qn('w:del'))
        del_elem.set(qn('w:id'), str(rev_id)); rev_id += 1
        del_elem.set(qn('w:author'), AUTHORS[2])
        del_elem.set(qn('w:date'), BASE_DATES[5])
        r = make_del_run_element(' with planned maintenance windows on Sundays')
        del_elem.append(r)

    # Change 7: Sarah Chen - Insert architecture detail
    if len(all_paras) > 22:
        p = all_paras[22]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[0])
        ins.set(qn('w:date'), BASE_DATES[6])
        r = make_run_element(' Service mesh implemented via Istio for inter-service communication.')
        ins.append(r)

    # Change 8: Marcus Rivera - Insert in timeline
    if len(all_paras) > 25:
        p = all_paras[min(25, len(all_paras) - 1)]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[1])
        ins.set(qn('w:date'), BASE_DATES[7])
        r = make_run_element(' Beta testing with select enterprise partners begins May 2026.')
        ins.append(r)

    # Change 9: Elena Kowalski - Delete text
    if len(all_paras) > 16:
        p = all_paras[16]
        del_elem = etree.SubElement(p, qn('w:del'))
        del_elem.set(qn('w:id'), str(rev_id)); rev_id += 1
        del_elem.set(qn('w:author'), AUTHORS[2])
        del_elem.set(qn('w:date'), BASE_DATES[8])
        r = make_del_run_element(' with a maximum of 10 concurrent waves')
        del_elem.append(r)

    # Change 10: Sarah Chen - Insert text
    if len(all_paras) > 17:
        p = all_paras[17]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[0])
        ins.set(qn('w:date'), BASE_DATES[9])
        r = make_run_element(' Reserved instance pricing optimization is included in cost projections.')
        ins.append(r)

    # Change 11: Marcus Rivera - Insert text about monitoring
    if len(all_paras) > 12:
        p = all_paras[12]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[1])
        ins.set(qn('w:date'), BASE_DATES[10])
        r = make_run_element(' Real-time alerting via PagerDuty integration is supported.')
        ins.append(r)

    # Change 12: Elena Kowalski - Delete old text
    if len(all_paras) > 10:
        p = all_paras[10]
        del_elem = etree.SubElement(p, qn('w:del'))
        del_elem.set(qn('w:id'), str(rev_id)); rev_id += 1
        del_elem.set(qn('w:author'), AUTHORS[2])
        del_elem.set(qn('w:date'), BASE_DATES[11])
        r = make_del_run_element(' including deprecated SNMP-based tools')
        del_elem.append(r)

    # Change 13: Sarah Chen - Insert compliance requirement
    if len(all_paras) > 20:
        p = all_paras[20]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[0])
        ins.set(qn('w:date'), BASE_DATES[12])
        r = make_run_element(' SOC 2 Type II compliance audit logging is mandatory.')
        ins.append(r)

    # Change 14: Marcus Rivera - Delete reference
    if len(all_paras) > 15:
        p = all_paras[15]
        del_elem = etree.SubElement(p, qn('w:del'))
        del_elem.set(qn('w:id'), str(rev_id)); rev_id += 1
        del_elem.set(qn('w:author'), AUTHORS[1])
        del_elem.set(qn('w:date'), BASE_DATES[13])
        r = make_del_run_element(' as specified in the original RFP document')
        del_elem.append(r)

    # Change 15: Elena Kowalski - Insert testing requirement
    if len(all_paras) > 18:
        p = all_paras[18]
        ins = etree.SubElement(p, qn('w:ins'))
        ins.set(qn('w:id'), str(rev_id)); rev_id += 1
        ins.set(qn('w:author'), AUTHORS[2])
        ins.set(qn('w:date'), BASE_DATES[14])
        r = make_run_element(' Automated chaos engineering tests shall validate rollback procedures weekly.')
        ins.append(r)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count tracked changes for verification
    doc2 = DocxDocument(OUTPUT)
    body2 = doc2.element.body
    ins_count = len(body2.findall('.//' + qn('w:ins')))
    del_count = len(body2.findall('.//' + qn('w:del')))
    print(f'Tracked changes: {ins_count} insertions, {del_count} deletions, total: {ins_count + del_count}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
