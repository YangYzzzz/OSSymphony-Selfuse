"""
Initial Setup: Mail merge template letter with placeholders
Task ID: writer_legal_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Firm letterhead ---
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run("MORRISON, CHEN & ASSOCIATES")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = subtitle.add_run("Attorneys at Law")
    run_sub.font.size = Pt(12)
    run_sub.font.name = "Times New Roman"
    run_sub.italic = True

    address_line = doc.add_paragraph()
    address_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_addr = address_line.add_run(
        "1200 Pacific Avenue, Suite 450  |  San Francisco, CA 94115\n"
        "Tel: (415) 555-0187  |  Fax: (415) 555-0188  |  info@morrisonchen.law"
    )
    run_addr.font.size = Pt(9)
    run_addr.font.name = "Times New Roman"
    run_addr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Horizontal rule (simulated with bottom border) ---
    doc.add_paragraph()  # spacer

    # --- Date ---
    date_para = doc.add_paragraph()
    date_run = date_para.add_run("April 2, 2026")
    date_run.font.size = Pt(11)
    date_run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # --- Recipient address block (with placeholders) ---
    addr_para = doc.add_paragraph()
    addr_run = addr_para.add_run("[CLIENT_NAME]")
    addr_run.font.size = Pt(11)
    addr_run.font.name = "Times New Roman"

    addr_para2 = doc.add_paragraph()
    addr_run2 = addr_para2.add_run("[CLIENT_ADDRESS]")
    addr_run2.font.size = Pt(11)
    addr_run2.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # --- Subject line ---
    subj_para = doc.add_paragraph()
    subj_run = subj_para.add_run("Re: Case No. [CASE_NUMBER] — Hearing Notification")
    subj_run.bold = True
    subj_run.font.size = Pt(11)
    subj_run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # --- Greeting ---
    greet = doc.add_paragraph()
    greet_run = greet.add_run("Dear [CLIENT_NAME],")
    greet_run.font.size = Pt(11)
    greet_run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # --- Body paragraph 1 ---
    body1 = doc.add_paragraph()
    body1_run = body1.add_run(
        "We are writing to inform you that a hearing has been scheduled in connection "
        "with your case, Case No. [CASE_NUMBER]. This hearing is an important step in "
        "the legal proceedings, and your attendance is strongly recommended."
    )
    body1_run.font.size = Pt(11)
    body1_run.font.name = "Times New Roman"
    body1.paragraph_format.space_after = Pt(6)

    # --- Body paragraph 2 ---
    body2 = doc.add_paragraph()
    body2_run = body2.add_run(
        "The hearing is scheduled for [HEARING_DATE] at the San Francisco County "
        "Superior Court, Department 302, located at 400 McAllister Street, San Francisco, "
        "CA 94102. Please plan to arrive at least thirty (30) minutes prior to the "
        "scheduled time to allow for security screening and check-in procedures."
    )
    body2_run.font.size = Pt(11)
    body2_run.font.name = "Times New Roman"
    body2.paragraph_format.space_after = Pt(6)

    # --- Body paragraph 3 ---
    body3 = doc.add_paragraph()
    body3_run = body3.add_run(
        "Please bring a valid government-issued photo identification, any documents or "
        "evidence previously discussed with your legal team, and any correspondence you "
        "have received from the court. If you are unable to attend on the scheduled date, "
        "please contact our office immediately so that we may request a continuance on your "
        "behalf."
    )
    body3_run.font.size = Pt(11)
    body3_run.font.name = "Times New Roman"
    body3.paragraph_format.space_after = Pt(6)

    # --- Body paragraph 4 ---
    body4 = doc.add_paragraph()
    body4_run = body4.add_run(
        "Should you have any questions regarding this matter or need to discuss your case "
        "further, please do not hesitate to reach out to our office at (415) 555-0187 or "
        "via email at hearings@morrisonchen.law. We remain committed to providing you with "
        "diligent and effective legal representation."
    )
    body4_run.font.size = Pt(11)
    body4_run.font.name = "Times New Roman"
    body4.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()  # spacer

    # --- Closing ---
    closing = doc.add_paragraph()
    closing_run = closing.add_run("Sincerely,")
    closing_run.font.size = Pt(11)
    closing_run.font.name = "Times New Roman"

    doc.add_paragraph()  # signature space
    doc.add_paragraph()  # signature space

    sig = doc.add_paragraph()
    sig_run = sig.add_run("Victoria A. Morrison, Esq.")
    sig_run.bold = True
    sig_run.font.size = Pt(11)
    sig_run.font.name = "Times New Roman"

    title = doc.add_paragraph()
    title_run = title.add_run("Senior Partner\nMorrison, Chen & Associates")
    title_run.font.size = Pt(11)
    title_run.font.name = "Times New Roman"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
