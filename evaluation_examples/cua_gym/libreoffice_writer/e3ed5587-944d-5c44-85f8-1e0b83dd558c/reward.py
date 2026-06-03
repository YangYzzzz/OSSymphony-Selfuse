"""
FINAL REWARD SCRIPT - SUCCESS
Task: Quick LibreOffice tweak: how can I set the line spacing of just the very last paragraph to Proportional 120%?
Generated: 2025-09-10 15:31:09
Status: success
Model: azure-o3
Total Steps: 15
"""

import os
from docx import Document

# ---------------------------------------------------------
# Reward Script :  Verify only the very last paragraph has
#                  line-spacing set to Proportional 120 %.
# ---------------------------------------------------------
# Scoring (progressive)
#   0.7  – last paragraph is set to 120 %   (≈ 288 twips)
#   0.3  – no earlier paragraph is 120 %
#   1.0  – both conditions satisfied
# ---------------------------------------------------------

FILE_PATH = (
    "/home/user/quick_libreoffice_tweak_how_can_i_set_the_line_spacing_of_just_the_very_last_paragraph_to_proportion.docx"
)

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _spacing(paragraph):
    """Return (lineRule, lineTwips) tuple or (None, None) if not present."""
    el = paragraph._p.find(".//w:pPr/w:spacing", NS)
    if el is None:
        return None, None
    return (
        el.get(f"{{{NS['w']}}}lineRule"),
        el.get(f"{{{NS['w']}}}line"),
    )


def verify_task(file_path: str = FILE_PATH) -> float:
    score = 0.0

    # ---------- prerequisite: file must exist & load ----------
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0
    try:
        doc = Document(file_path)
        print(f"✓ File loaded – {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"✗ Unable to open DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- locate last non-empty paragraph ----------
    last_idx = next(
        (i for i in range(len(doc.paragraphs) - 1, -1, -1) if doc.paragraphs[i].text.strip()),
        None,
    )
    if last_idx is None:
        print("✗ Document contains no non-empty paragraphs")
        print("REWARD: 0.0")
        return 0.0
    print(f"Last non-empty paragraph index: {last_idx}")

    # ---------- requirement 1 – correct spacing on last paragraph ----------
    lr_last, ln_last = _spacing(doc.paragraphs[last_idx])
    correct_last = False
    try:
        if lr_last == "auto" and ln_last is not None and abs(int(ln_last) - 288) <= 5:
            correct_last = True
    except ValueError:
        pass
    if correct_last:
        print("✓ Last paragraph set to Proportional 120 % (≈288 twips)")
        score += 0.7
    else:
        print("✗ Last paragraph is NOT set to Proportional 120 %")

    # ---------- requirement 2 – nobody else has that spacing ----------
    others_ok = True
    for i, para in enumerate(doc.paragraphs):
        if i == last_idx:
            continue
        lr_i, ln_i = _spacing(para)
        try:
            if lr_i == "auto" and ln_i is not None and abs(int(ln_i) - 288) <= 5:
                print(f"✗ Paragraph {i} also has 120 % spacing – should not")
                others_ok = False
                break
        except ValueError:
            continue
    if others_ok:
        print("✓ All other paragraphs keep their original spacing")
        score += 0.3
    else:
        print("✗ Additional paragraphs share 120 % spacing")

    final_score = min(score, 1.0)
    print(f"REWARD: {final_score}")
    return final_score


if __name__ == "__main__":
    verify_task(FILE_PATH)
