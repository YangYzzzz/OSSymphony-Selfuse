"""
Reward Script: Color vowel-starting words red and consonant-starting words blue in introduction paragraph.
Task ID: osworld_writer_vowel_consonant_coloring_005
Domain: libreoffice_writer
Scoring:
  Component 1: Words in intro paragraph are individually colored (not a single uncolored run) — 0.4 pts
  Component 2: All vowel-starting words colored red (FF0000) — 0.35 pts
  Component 3: All consonant-starting words colored blue (0000FF) — 0.25 pts
"""

import os

from docx import Document
from docx.shared import RGBColor

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_005'

VOWELS = set('aeiouAEIOU')

RED = RGBColor(0xFF, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xFF)

COLOR_TOLERANCE = 30  # Euclidean RGB distance tolerance


def color_distance(c1, c2):
    """Euclidean distance between two RGBColor objects."""
    return ((c1[0] - c2[0]) ** 2 + (c1[1] - c2[1]) ** 2 + (c1[2] - c2[2]) ** 2) ** 0.5


def is_red(rgb):
    return color_distance(rgb, RED) < COLOR_TOLERANCE


def is_blue(rgb):
    return color_distance(rgb, BLUE) < COLOR_TOLERANCE


def find_intro_paragraph(doc):
    """
    Find the introduction paragraph: the Normal-style paragraph immediately
    after the 'Introduction' heading.
    """
    heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') and 'introduction' in para.text.lower():
            heading_idx = i
            break
    if heading_idx is not None:
        for para in doc.paragraphs[heading_idx + 1:]:
            if para.style.name in ('Normal', 'Body Text'):
                return para
    # Fallback: return the first Normal paragraph that looks like intro text
    for para in doc.paragraphs:
        if para.style.name in ('Normal', 'Body Text') and len(para.text) > 50:
            return para
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the introduction paragraph
    intro_para = find_intro_paragraph(doc)
    if intro_para is None:
        print("CRITICAL: Could not find introduction paragraph")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Introduction paragraph text (first 80 chars): {intro_para.text[:80]!r}")
    print(f"INFO: Number of runs in intro paragraph: {len(intro_para.runs)}")

    # Component 1: Words in intro paragraph are individually colored (0.4 points)
    # In the initial state, there is a single run with no color applied.
    # In the golden state, each word is its own colored run (many runs, each with a color).
    try:
        word_runs = [r for r in intro_para.runs if r.text.strip()]  # runs with actual word text
        colored_word_runs = [
            r for r in word_runs
            if r.font.color.type is not None and r.font.color.rgb is not None
        ]

        # Expect that most (>=90%) of word-bearing runs have a color applied
        if len(word_runs) == 0:
            print("FAIL: Component 1 — No word runs found in introduction paragraph")
        else:
            ratio = len(colored_word_runs) / len(word_runs)
            if ratio >= 0.9:
                print(f"PASS: Component 1 — {len(colored_word_runs)}/{len(word_runs)} word runs are colored ({ratio:.0%}), satisfies coloring requirement (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 — Only {len(colored_word_runs)}/{len(word_runs)} word runs are colored ({ratio:.0%}); expected >=90%")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All vowel-starting words colored red FF0000 (0.35 points)
    try:
        vowel_total = 0
        vowel_correct = 0
        vowel_errors = []

        for run in intro_para.runs:
            word = run.text.strip().strip('.,!?;:"\'')
            if not word:
                continue  # skip spaces and punctuation-only runs
            first_char = word[0]
            if first_char in VOWELS:
                vowel_total += 1
                if run.font.color.type is not None and run.font.color.rgb is not None:
                    if is_red(run.font.color.rgb):
                        vowel_correct += 1
                    else:
                        vowel_errors.append(f"  '{run.text}' has color {run.font.color.rgb} (expected red FF0000)")
                else:
                    vowel_errors.append(f"  '{run.text}' has no color (expected red FF0000)")

        if vowel_total == 0:
            print("FAIL: Component 2 — No vowel-starting words found")
        else:
            ratio = vowel_correct / vowel_total
            if ratio >= 0.9:
                print(f"PASS: Component 2 — {vowel_correct}/{vowel_total} vowel-starting words correctly colored red ({ratio:.0%}) (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 2 — Only {vowel_correct}/{vowel_total} vowel-starting words colored red ({ratio:.0%})")
                for err in vowel_errors[:5]:
                    print(err)
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All consonant-starting words colored blue 0000FF (0.25 points)
    try:
        consonant_total = 0
        consonant_correct = 0
        consonant_errors = []

        for run in intro_para.runs:
            word = run.text.strip().strip('.,!?;:"\'')
            if not word:
                continue  # skip spaces and punctuation-only runs
            first_char = word[0]
            if first_char.isalpha() and first_char not in VOWELS:
                consonant_total += 1
                if run.font.color.type is not None and run.font.color.rgb is not None:
                    if is_blue(run.font.color.rgb):
                        consonant_correct += 1
                    else:
                        consonant_errors.append(f"  '{run.text}' has color {run.font.color.rgb} (expected blue 0000FF)")
                else:
                    consonant_errors.append(f"  '{run.text}' has no color (expected blue 0000FF)")

        if consonant_total == 0:
            print("FAIL: Component 3 — No consonant-starting words found")
        else:
            ratio = consonant_correct / consonant_total
            if ratio >= 0.9:
                print(f"PASS: Component 3 — {consonant_correct}/{consonant_total} consonant-starting words correctly colored blue ({ratio:.0%}) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — Only {consonant_correct}/{consonant_total} consonant-starting words colored blue ({ratio:.0%})")
                for err in consonant_errors[:5]:
                    print(err)
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
