"""
Initial Setup: Vowel/Consonant Word Coloring in Technical Report Introduction
Task ID: osworld_writer_vowel_consonant_coloring_005
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_vowel_consonant_coloring_005'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Enterprise AI Adoption: A Technical Overview', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction section heading
    doc.add_heading('Introduction', level=1)

    # Introduction paragraph — 3 sentences, ~45 words, NO font colors applied
    intro_para = doc.add_paragraph(
        'Artificial intelligence has emerged as a transformative technology across modern industries '
        'over recent years. '
        'Organizations worldwide are investing heavily in automated systems to optimize operational '
        'efficiency and significantly reduce costs. '
        'This report analyzes current trends, key implementation challenges, and strategic outcomes '
        'observed from large enterprise deployments.'
    )

    # Section 2: Background
    doc.add_heading('Background', level=1)
    bg_para = doc.add_paragraph(
        'The rapid advancement of machine learning and neural network architectures over the past decade '
        'has enabled capabilities that were previously considered theoretical. Starting from narrow '
        'task-specific models, modern AI systems now demonstrate generalization across diverse problem domains. '
        'Research institutions such as Stanford University and MIT have published extensive studies documenting '
        'performance benchmarks that surpass human-level accuracy on structured tasks.'
    )

    # Section 3: Methodology
    doc.add_heading('Methodology', level=1)
    method_para = doc.add_paragraph(
        'Data collection was conducted across fourteen enterprise organizations between January 2024 and '
        'September 2024. Interviews were performed with senior technology officers, department heads, and '
        'operational staff to capture diverse perspectives on adoption barriers and success factors. '
        'Quantitative metrics including deployment timelines, cost reduction percentages, and employee '
        'productivity scores were gathered through structured surveys and system logs.'
    )

    # Section 4: Results
    doc.add_heading('Results', level=1)
    results_para = doc.add_paragraph(
        'Preliminary findings indicate that organizations deploying AI-assisted workflows reported an '
        'average productivity increase of 34% within the first six months of implementation. '
        'Customer satisfaction scores improved by 22 percentage points on average, while operational '
        'costs decreased by approximately 18% across all participating organizations. '
        'Error rates in data processing tasks dropped significantly from 8.3% to under 1.2% following '
        'full system integration.'
    )

    # Section 5: Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion_para = doc.add_paragraph(
        'The evidence presented in this report strongly supports continued investment in enterprise AI '
        'adoption strategies. Organizations that approach implementation with clear objectives, adequate '
        'training resources, and robust governance frameworks consistently achieve superior outcomes. '
        'Future research should focus on long-term sustainability, workforce adaptation, and ethical '
        'considerations associated with large-scale automation initiatives.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
