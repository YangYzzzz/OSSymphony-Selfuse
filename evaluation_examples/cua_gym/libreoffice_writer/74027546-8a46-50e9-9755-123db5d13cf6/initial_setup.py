"""
Initial Setup: Novel manuscript without front matter pages
Task ID: writer_creative_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_053'
# The task context says file is on ~/Desktop/
OUTPUT = f'{WORKDIR}/Desktop/novel_manuscript.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Make sure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Set default style to 12pt, left-aligned (no front matter)
    # First line: Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_title = p_title.add_run('Echoes of Tomorrow')
    run_title.font.size = Pt(12)

    # Second line: Author
    p_author = doc.add_paragraph()
    p_author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_author = p_author.add_run('by Marcus Reeves')
    run_author.font.size = Pt(12)

    # 25 paragraphs of novel text starting with Chapter 1
    novel_paragraphs = [
        'Chapter 1',
        'The morning sun cast long shadows across the abandoned train station, its rusted tracks disappearing into the dense forest beyond.',
        'Elena Vasquez adjusted her backpack straps and checked her compass one last time. The coordinates she had received were precise: 47.3829° N, 19.1147° E.',
        'She had been chasing this lead for three years. Three years of dead ends, false trails, and near misses. But this time felt different.',
        'The station door groaned as she pushed it open, its hinges protesting against years of disuse. Inside, dust motes danced in the pale light filtering through broken windows.',
        'On the central bench sat a weathered leather satchel, exactly as the anonymous tip had promised.',
        'Elena approached cautiously, her eyes scanning the shadows for any sign of movement. The silence was absolute, broken only by the distant call of a crow.',
        'She knelt beside the satchel and carefully opened the brass clasps. Inside were papers, dozens of them, covered in a neat, precise handwriting.',
        'The top document bore a date: March 15, 1987. Nearly forty years ago.',
        'Her breath caught as she read the first line: "Project Echo has been terminated. The subjects must never be found."',
        'Chapter 2',
        'Three weeks earlier, Elena had been sitting in her cramped office at the University of Budapest, grading undergraduate papers on Cold War espionage.',
        'The phone had rung at precisely 11:47 PM, an hour when only emergencies called.',
        '"Professor Vasquez," said the voice, low and urgent. "I have information about your father."',
        'Her father, Dr. Andrei Vasquez, had disappeared in 1989, two weeks before the fall of the Berlin Wall.',
        'She had spent her entire academic career studying the period, secretly hoping that somewhere in the historical record lay the answer to his disappearance.',
        '"Who is this?" she had demanded.',
        '"Someone who knew him. Someone who knows what they did to him."',
        'The line had gone dead before she could ask anything more.',
        'Elena had not slept that night. Instead, she had pulled out the worn photograph she always kept in her desk drawer.',
        'In it, her father stood before a university lecture hall, mid-sentence, his eyes alight with the passion that had defined him.',
        'She was seven years old in the photo, seated in the front row, watching him with undisguised admiration.',
        'Now, standing in that desolate train station with his secrets in her hands, she felt that same mixture of awe and terror.',
        'The papers trembled slightly as she lifted them. Outside, a wind had picked up, rattling the broken windowpanes.',
        'Whatever lay hidden in these documents had been worth protecting for nearly four decades. She could only hope it would still be worth the cost of uncovering it.',
    ]

    for para_text in novel_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run(para_text)
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
