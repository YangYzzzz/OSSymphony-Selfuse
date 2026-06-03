"""
Initial Setup: Academic paper with References section (no hanging indent)
Task ID: wrpara_012
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_012'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ---- Title ----
    title = doc.add_heading('The Impact of Remote Work on Organizational Productivity: A Meta-Analysis', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Authors ----
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Elena Martinez, David Park, and Rajesh Gupta')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil.add_run('Department of Management Sciences, Northwestern University')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.name = 'Times New Roman'

    # ---- Abstract ----
    doc.add_heading('Abstract', level=2)
    abstract_text = (
        'This meta-analysis synthesizes findings from 47 empirical studies conducted between '
        '2019 and 2024 examining the relationship between remote work arrangements and '
        'organizational productivity. Using a random-effects model, we analyzed data from '
        '23,450 participants across 14 countries. Results indicate a modest positive effect '
        'of remote work on individual productivity (d = 0.31, 95% CI [0.18, 0.44]), with '
        'significant moderators including industry type, job autonomy, and technological '
        'infrastructure. Hybrid models (2-3 days remote) showed the strongest effects '
        '(d = 0.42). Implications for organizational policy and future research directions '
        'are discussed.'
    )
    p = doc.add_paragraph(abstract_text)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # ---- Introduction ----
    doc.add_heading('Introduction', level=2)
    intro_paras = [
        (
            'The rapid shift to remote work precipitated by the COVID-19 pandemic fundamentally '
            'altered how organizations conceptualize workplace productivity. Prior to 2020, '
            'approximately 5.7% of the U.S. workforce engaged in regular remote work (Bureau of '
            'Labor Statistics, 2019). By mid-2020, this figure exceeded 42%, prompting both '
            'practical and scholarly interest in the productivity implications of distributed work '
            'arrangements (Brynjolfsson et al., 2020).'
        ),
        (
            'Early theoretical frameworks suggested competing predictions. Agency theory predicted '
            'reduced productivity due to diminished monitoring (Jensen & Meckling, 1976), while '
            'self-determination theory anticipated gains from increased autonomy (Deci & Ryan, '
            '2000). Empirical evidence has been mixed, with studies reporting effects ranging from '
            'a 13% productivity increase in call centers (Bloom et al., 2015) to significant '
            'declines in collaborative innovation tasks (Yang et al., 2022).'
        ),
        (
            'This meta-analysis addresses the need for a comprehensive synthesis of the growing '
            'body of empirical research. We examine overall effect sizes, identify key moderators, '
            'and provide evidence-based recommendations for organizational decision-making regarding '
            'remote work policies.'
        ),
    ]
    for text in intro_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # ---- Method ----
    doc.add_heading('Method', level=2)
    method_paras = [
        (
            'We conducted a systematic literature search using PsycINFO, Web of Science, and '
            'Google Scholar databases. Search terms included combinations of "remote work," '
            '"telecommuting," "work from home," and "productivity." Initial screening yielded '
            '312 articles, of which 47 met inclusion criteria: (a) empirical study with '
            'quantifiable productivity measures, (b) comparison between remote and in-office '
            'conditions, and (c) published in peer-reviewed outlets between January 2019 and '
            'December 2024.'
        ),
        (
            'Effect sizes were computed as Cohen\'s d for each study. A random-effects model was '
            'employed to account for between-study heterogeneity. Moderator analyses were conducted '
            'using mixed-effects meta-regression. Publication bias was assessed via funnel plot '
            'asymmetry and Egger\'s regression test.'
        ),
    ]
    for text in method_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # ---- Results ----
    doc.add_heading('Results', level=2)
    results_paras = [
        (
            'The overall weighted mean effect size was d = 0.31 (95% CI [0.18, 0.44], p < .001), '
            'indicating a small-to-medium positive effect of remote work on productivity. '
            'Significant heterogeneity was observed (Q(46) = 187.3, p < .001; I-squared = 75.4%), '
            'warranting moderator analysis.'
        ),
        (
            'Industry type emerged as the strongest moderator (QB(3) = 24.7, p < .001). '
            'Knowledge-intensive industries (technology, finance, consulting) showed the largest '
            'effects (d = 0.48), while service industries showed negligible effects (d = 0.07). '
            'Hybrid arrangements of 2-3 remote days per week yielded higher effect sizes (d = 0.42) '
            'compared to fully remote configurations (d = 0.21).'
        ),
    ]
    for text in results_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # ---- Discussion ----
    doc.add_heading('Discussion', level=2)
    disc_text = (
        'Our findings provide nuanced support for the productivity benefits of remote work, '
        'while highlighting the importance of contextual factors. The superiority of hybrid '
        'models aligns with recent theorizing about the complementary roles of focused '
        'individual work and in-person collaboration (Choudhury et al., 2021). Organizations '
        'considering permanent remote work policies should attend to industry-specific factors, '
        'invest in digital infrastructure, and ensure adequate managerial training for '
        'distributed team leadership.'
    )
    p = doc.add_paragraph(disc_text)
    p.paragraph_format.first_line_indent = Cm(1.27)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # ---- References ---- (NO HANGING INDENT - this is what the task requires the agent to add)
    doc.add_heading('References', level=2)

    references = [
        'Bloom, N., Liang, J., Roberts, J., & Ying, Z. J. (2015). Does working from home work? Evidence from a Chinese experiment. The Quarterly Journal of Economics, 130(1), 165-218. https://doi.org/10.1093/qje/qju032',
        'Brynjolfsson, E., Horton, J. J., Ozimek, A., Rock, D., Sharma, G., & TuYe, H. Y. (2020). COVID-19 and remote work: An early look at US data. National Bureau of Economic Research Working Paper No. 27344.',
        'Choudhury, P., Foroughi, C., & Larson, B. (2021). Work-from-anywhere: The productivity effects of geographic flexibility. Strategic Management Journal, 42(4), 655-683. https://doi.org/10.1002/smj.3251',
        'Deci, E. L., & Ryan, R. M. (2000). The "what" and "why" of goal pursuits: Human needs and the self-determination of behavior. Psychological Inquiry, 11(4), 227-268. https://doi.org/10.1207/S15327965PLI1104_01',
        'Jensen, M. C., & Meckling, W. H. (1976). Theory of the firm: Managerial behavior, agency costs and ownership structure. Journal of Financial Economics, 3(4), 305-360. https://doi.org/10.1016/0304-405X(76)90026-X',
        'Yang, L., Holtz, D., Jaffe, S., Suri, S., Sinha, S., Weston, J., Joyce, C., Shah, N., Sherman, K., Hecht, B., & Teevan, J. (2022). The effects of remote work on collaboration among information workers. Nature Human Behaviour, 6(1), 43-54. https://doi.org/10.1038/s41562-021-01196-4',
    ]

    for ref_text in references:
        p = doc.add_paragraph(ref_text)
        # No indentation at all - plain paragraphs
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
