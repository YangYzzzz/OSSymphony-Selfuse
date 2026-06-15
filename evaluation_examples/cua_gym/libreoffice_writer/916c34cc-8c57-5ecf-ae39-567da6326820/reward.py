"""
Reward Script: Add hanging indent to bibliography paragraphs
Task ID: wrpara_012
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): All 6 reference paragraphs have left_indent ~1.27cm
  Component 2 (0.5): All 6 reference paragraphs have first_line_indent ~-1.27cm (hanging)
"""

import os
from docx import Document
from docx.shared import Emu, Cm

WORKDIR = '/home/user'
TASK_ID = 'wrpara_012'

# Tolerance: 1.27cm = 0.5 inches. Allow +/-0.15cm tolerance for rounding
EXPECTED_LEFT_INDENT_CM = 1.27
EXPECTED_FIRST_LINE_INDENT_CM = -1.27
TOLERANCE_CM = 0.15


def find_reference_paragraphs(doc):
    """
    Find the 'References' heading and return the paragraph indices
    of the bibliography entries that follow it.
    """
    ref_heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name.startswith('Heading') and para.text.strip().lower() == 'references':
            ref_heading_idx = i
            break

    if ref_heading_idx is None:
        return None, []

    # Collect all Normal-style paragraphs after the References heading
    bib_indices = []
    for j in range(ref_heading_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[j]
        # Stop if we hit another heading
        if para.style and para.style.name.startswith('Heading'):
            break
        # Only count non-empty paragraphs
        if para.text.strip():
            bib_indices.append(j)

    return ref_heading_idx, bib_indices


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the reference paragraphs
    ref_heading_idx, bib_indices = find_reference_paragraphs(doc)

    if ref_heading_idx is None:
        print("CRITICAL: No 'References' heading found in document")
        print("REWARD: 0.0")
        return 0.0

    if len(bib_indices) == 0:
        print("CRITICAL: No bibliography paragraphs found after References heading")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found {len(bib_indices)} bibliography paragraphs after 'References' heading (para {ref_heading_idx})")

    # Component 1: All bibliography paragraphs have left_indent ~1.27cm (0.5 points)
    try:
        left_indent_pass_count = 0
        for idx in bib_indices:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            left_indent = pf.left_indent
            if left_indent is not None:
                li_cm = Emu(left_indent).cm
                if abs(li_cm - EXPECTED_LEFT_INDENT_CM) <= TOLERANCE_CM:
                    left_indent_pass_count += 1
                    print(f"  PASS: Para {idx} left_indent={li_cm:.4f}cm (expected ~{EXPECTED_LEFT_INDENT_CM}cm)")
                else:
                    print(f"  FAIL: Para {idx} left_indent={li_cm:.4f}cm (expected ~{EXPECTED_LEFT_INDENT_CM}cm)")
            else:
                print(f"  FAIL: Para {idx} left_indent=None (expected ~{EXPECTED_LEFT_INDENT_CM}cm)")

        if left_indent_pass_count == len(bib_indices):
            print(f"PASS: Component 1 - All {len(bib_indices)} bib paragraphs have correct left_indent (0.5 pts)")
            total_score += 0.5
        elif left_indent_pass_count > 0:
            partial = 0.5 * (left_indent_pass_count / len(bib_indices))
            print(f"PARTIAL: Component 1 - {left_indent_pass_count}/{len(bib_indices)} bib paragraphs have correct left_indent ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - No bib paragraphs have correct left_indent")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: All bibliography paragraphs have first_line_indent ~-1.27cm (0.5 points)
    try:
        fli_pass_count = 0
        for idx in bib_indices:
            para = doc.paragraphs[idx]
            pf = para.paragraph_format
            first_line_indent = pf.first_line_indent
            if first_line_indent is not None:
                fli_cm = Emu(first_line_indent).cm
                if abs(fli_cm - EXPECTED_FIRST_LINE_INDENT_CM) <= TOLERANCE_CM:
                    fli_pass_count += 1
                    print(f"  PASS: Para {idx} first_line_indent={fli_cm:.4f}cm (expected ~{EXPECTED_FIRST_LINE_INDENT_CM}cm)")
                else:
                    print(f"  FAIL: Para {idx} first_line_indent={fli_cm:.4f}cm (expected ~{EXPECTED_FIRST_LINE_INDENT_CM}cm)")
            else:
                print(f"  FAIL: Para {idx} first_line_indent=None (expected ~{EXPECTED_FIRST_LINE_INDENT_CM}cm)")

        if fli_pass_count == len(bib_indices):
            print(f"PASS: Component 2 - All {len(bib_indices)} bib paragraphs have correct first_line_indent (0.5 pts)")
            total_score += 0.5
        elif fli_pass_count > 0:
            partial = 0.5 * (fli_pass_count / len(bib_indices))
            print(f"PARTIAL: Component 2 - {fli_pass_count}/{len(bib_indices)} bib paragraphs have correct first_line_indent ({partial:.2f} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - No bib paragraphs have correct first_line_indent")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
