"""
Reward Script: HR Department Newsletter with Complex Layout
Task ID: writer_hr_065
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20) - Table-based layout structure (>= 3 tables vs initial 0)
  Component 2 (0.20) - Banner header with company name, newsletter title, date (centered, styled)
  Component 3 (0.20) - Two-column article section with 3 articles
  Component 4 (0.20) - Sidebar with quick stats (headcount, positions, turnover rate)
  Component 5 (0.20) - Pull quote with italic styling + footer with contact info
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_065'


def persist_app_state(domain):
    """Best-effort save via Ctrl+S in case document is open in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # =========================================================================
    # Component 1: Table-based layout structure (0.20 points)
    # Initial has 0 tables. Golden uses 3 tables for layout.
    # We require at least 2 tables (layout tables for banner, content, footer).
    # =========================================================================
    try:
        num_tables = len(doc.tables)
        if num_tables >= 3:
            print(f"PASS: Component 1 — {num_tables} tables found (>= 3 required) (0.20 pts)")
            total_score += 0.20
        elif num_tables >= 2:
            print(f"PARTIAL: Component 1 — {num_tables} tables found (>= 3 ideal, >= 2 partial) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 1 — Only {num_tables} tables found, need >= 2 for layout")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Banner header with company name, newsletter title, date (0.20 points)
    # Golden: Table 0 is a single-cell table with centered text containing
    # "MERIDIAN TECHNOLOGIES" and "Newsletter" and a date reference.
    # Initial: no tables, just plain text paragraphs.
    # =========================================================================
    try:
        banner_found = False
        if num_tables >= 1:
            t0 = doc.tables[0]
            # Get all text from first table
            banner_text = ""
            for row in t0.rows:
                for cell in row.cells:
                    banner_text += cell.text.lower() + " "

            has_company = "meridian" in banner_text
            has_newsletter = "newsletter" in banner_text
            has_date = "2026" in banner_text or "april" in banner_text

            # Check for centered alignment in banner
            is_centered = False
            for row in t0.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.paragraph_format.alignment is not None:
                            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                            if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                                is_centered = True
                                break

            # Check for styling (bold or colored text)
            has_styling = False
            for row in t0.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.bold or (r.font.color and r.font.color.rgb):
                                has_styling = True
                                break

            if has_company and has_newsletter and has_date and is_centered and has_styling:
                print(f"PASS: Component 2 — Banner header found with company, newsletter, date, centered, styled (0.20 pts)")
                total_score += 0.20
                banner_found = True
            elif has_company and has_newsletter:
                print(f"PARTIAL: Component 2 — Banner has company+newsletter but missing some elements (centered={is_centered}, date={has_date}, styled={has_styling}) (0.10 pts)")
                total_score += 0.10
                banner_found = True
            else:
                print(f"FAIL: Component 2 — Banner missing key content (company={has_company}, newsletter={has_newsletter})")
        else:
            print(f"FAIL: Component 2 — No tables found for banner header")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: Two-column article section with 3 articles (0.20 points)
    # Golden: Table 1 has 2 columns. Left column contains article text with
    # "New Hires", "Policy Update"/"Flexible Work", and "Upcoming Events".
    # Initial: plain text paragraphs, no table structure.
    # =========================================================================
    try:
        two_col_found = False
        articles_found = 0

        # Search through all tables for a two-column layout with articles
        for table in doc.tables:
            if len(table.columns) >= 2:
                two_col_found = True
                # Check left column for articles
                left_text = ""
                for row in table.rows:
                    left_text += row.cells[0].text.lower() + " "

                if "new hire" in left_text or "welcome" in left_text:
                    articles_found += 1
                if "policy" in left_text or "flexible work" in left_text:
                    articles_found += 1
                if "upcoming event" in left_text or "events" in left_text:
                    articles_found += 1
                break

        if two_col_found and articles_found >= 3:
            print(f"PASS: Component 3 — Two-column layout with {articles_found} articles found (0.20 pts)")
            total_score += 0.20
        elif two_col_found and articles_found >= 2:
            print(f"PARTIAL: Component 3 — Two-column layout with {articles_found}/3 articles (0.10 pts)")
            total_score += 0.10
        elif two_col_found:
            print(f"PARTIAL: Component 3 — Two-column layout found but only {articles_found}/3 articles (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 3 — No two-column table layout found (two_col={two_col_found}, articles={articles_found})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================================
    # Component 4: Sidebar with quick stats (0.20 points)
    # Golden: Right column of Table 1 has stats: headcount (1,247), open
    # positions (34), turnover rate (8.3%), with styled numbers.
    # Initial: stats are just plain paragraphs, not in a sidebar.
    # =========================================================================
    try:
        sidebar_found = False
        stats_count = 0

        for table in doc.tables:
            if len(table.columns) >= 2:
                # Check right column for stats
                right_text = ""
                for row in table.rows:
                    right_text += row.cells[-1].text.lower() + " "

                # Check for stat keywords
                if "1,247" in right_text or "1247" in right_text or "headcount" in right_text:
                    stats_count += 1
                if "34" in right_text and ("position" in right_text or "open" in right_text):
                    stats_count += 1
                if "8.3" in right_text or "turnover" in right_text:
                    stats_count += 1

                if stats_count >= 2:
                    sidebar_found = True
                    break

        if sidebar_found and stats_count >= 3:
            print(f"PASS: Component 4 — Sidebar with {stats_count} quick stats found (0.20 pts)")
            total_score += 0.20
        elif sidebar_found:
            print(f"PARTIAL: Component 4 — Sidebar found with {stats_count}/3 stats (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — No sidebar with quick stats found (stats_count={stats_count})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =========================================================================
    # Component 5: Pull quote with italic styling + Footer with contact info (0.20 points)
    # Golden: Pull quote in italic in sidebar. Footer table with contact info.
    # Initial: quotes are plain text paragraphs, no footer table.
    # Sub-components: pull quote (0.10) + footer (0.10)
    # =========================================================================
    try:
        # Check for pull quote: italic text containing a quote with attribution
        pull_quote_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text.strip()
                        # Look for quoted text (with typographic or straight quotes)
                        is_quote = (
                            (text.startswith('"') or text.startswith('\u201c'))
                            and len(text) > 20
                        )
                        if is_quote:
                            # Check if runs are italic
                            for r in p.runs:
                                if r.italic:
                                    pull_quote_found = True
                                    break
                        # Also check for attribution lines like "- Henry Ford"
                        if text.startswith('—') or text.startswith('-'):
                            for r in p.runs:
                                if r.italic:
                                    pull_quote_found = True
                                    break

        if pull_quote_found:
            print(f"PASS: Component 5a — Pull quote with italic styling found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5a — No pull quote with italic styling found in tables")

        # Check for footer with contact info in a table
        footer_found = False
        for table in doc.tables:
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text.lower() + " "
            if "contact" in table_text and ("email" in table_text or "phone" in table_text):
                # Verify it has actual contact details
                if "hr@" in table_text or "555" in table_text or "meridiantech" in table_text:
                    footer_found = True
                    break

        if footer_found:
            print(f"PASS: Component 5b — Footer table with contact info found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5b — No footer table with contact info found")

    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
