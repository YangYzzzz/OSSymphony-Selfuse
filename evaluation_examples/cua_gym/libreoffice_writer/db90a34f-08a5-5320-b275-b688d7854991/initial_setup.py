"""
Initial Setup: HR Department Newsletter - Raw content in plain text
Task ID: writer_hr_065
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_065'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Raw plain text content - no formatting, no layout structures
    # Just the content that needs to be organized into a newsletter

    doc.add_paragraph("Meridian Technologies Inc.")
    doc.add_paragraph("HR Department Newsletter")
    doc.add_paragraph("April 2026 Edition")
    doc.add_paragraph("")

    doc.add_paragraph("New Hires Announcement")
    doc.add_paragraph(
        "We are thrilled to welcome several talented individuals to our growing team this month. "
        "Please join us in giving a warm welcome to each of our newest colleagues as they begin "
        "their journeys with Meridian Technologies."
    )
    doc.add_paragraph(
        "Elena Vasquez joins the Software Engineering division as a Senior Backend Developer. "
        "Elena brings over eight years of experience in distributed systems and cloud architecture "
        "from her previous role at Cloudbridge Solutions. She holds a Master's degree in Computer "
        "Science from Stanford University and is passionate about building scalable microservices."
    )
    doc.add_paragraph(
        "David Okonkwo has been appointed as our new Regional Sales Manager for the Northeast territory. "
        "With a proven track record of exceeding quarterly targets by an average of 23% at his "
        "previous company, NovaTech Partners, David will be instrumental in expanding our enterprise "
        "client portfolio. He is based in our Boston office."
    )
    doc.add_paragraph(
        "Priya Sharma starts as a UX Research Lead in the Product Design team. Priya previously "
        "led user research initiatives at DesignForward Labs, where she developed the award-winning "
        "accessibility framework now adopted by over 200 organizations worldwide."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Policy Update: Flexible Work Arrangements")
    doc.add_paragraph(
        "Effective May 1, 2026, the company is introducing an updated Flexible Work Arrangement policy. "
        "After extensive feedback from our annual employee satisfaction survey, which received a record "
        "89% participation rate, management has approved the following key changes to support work-life balance."
    )
    doc.add_paragraph(
        "Employees may now choose to work remotely up to three days per week, an increase from the "
        "previous two-day limit. Core collaboration hours are established between 10:00 AM and 2:00 PM "
        "Eastern Time on all workdays to ensure team availability for meetings and cross-functional projects."
    )
    doc.add_paragraph(
        "A new compressed workweek option allows eligible employees to complete their 40-hour week "
        "in four 10-hour days, subject to manager approval. Additionally, the home office stipend "
        "has been increased from $500 to $750 annually to help employees maintain productive remote "
        "work environments. Please review the full policy document on the HR portal."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Upcoming Events")
    doc.add_paragraph(
        "April 15 - Spring Wellness Fair: Join us in the main lobby from 11 AM to 3 PM for health "
        "screenings, fitness demonstrations, nutrition consultations, and mental health resource booths. "
        "Free smoothies and healthy snacks will be provided courtesy of the Wellness Committee."
    )
    doc.add_paragraph(
        "April 22 - Earth Day Volunteer Day: Meridian Technologies is partnering with GreenCity "
        "Foundation for a company-wide volunteer event at Riverside Park. Activities include tree "
        "planting, trail restoration, and community garden setup. Transportation provided from all "
        "office locations. Sign up on the intranet by April 18."
    )
    doc.add_paragraph(
        "May 5-9 - Annual Leadership Summit: This year's summit theme is 'Innovation Through "
        "Inclusion.' Featured speakers include Dr. Amara Obi from the Global Diversity Institute "
        "and tech entrepreneur Marcus Lee. Sessions will cover inclusive hiring practices, "
        "neurodiversity in the workplace, and building psychologically safe teams."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Quick Stats")
    doc.add_paragraph("Total Headcount: 1,247")
    doc.add_paragraph("Open Positions: 34")
    doc.add_paragraph("Employee Turnover Rate: 8.3%")
    doc.add_paragraph("Average Tenure: 4.2 years")
    doc.add_paragraph("Training Hours (Q1): 12,580")
    doc.add_paragraph("Internal Promotions (Q1): 28")
    doc.add_paragraph("")

    doc.add_paragraph(
        "\"The strength of the team is each individual member. The strength of each member is the team.\" "
        "- Phil Jackson"
    )
    doc.add_paragraph("")

    doc.add_paragraph(
        "\"Coming together is a beginning, staying together is progress, and working together is success.\" "
        "- Henry Ford"
    )
    doc.add_paragraph("")

    doc.add_paragraph("Contact HR Department")
    doc.add_paragraph("Email: hr@meridiantech.com")
    doc.add_paragraph("Phone: (555) 234-5678 ext. 100")
    doc.add_paragraph("Office: Building A, Suite 302")
    doc.add_paragraph("HR Portal: https://hr.meridiantech.com")
    doc.add_paragraph("Office Hours: Monday-Friday, 8:00 AM - 5:00 PM ET")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
