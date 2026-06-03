"""
Initial Setup: Set line spacing for poetry stanzas
Task ID: wrpara_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'wrpara_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title paragraph
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('The Road Not Taken')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Times New Roman'
    # Default spacing - no special spacing set

    # The poem stanzas - each stanza is ONE paragraph with soft line breaks
    stanza1_lines = [
        'Two roads diverged in a yellow wood,',
        'And sorry I could not travel both',
        'And be one traveler, long I stood',
        'And looked down one as far as I could',
        'To where it bent in the undergrowth;',
    ]

    stanza2_lines = [
        'Then took the other, as just as fair,',
        'And having perhaps the better claim,',
        'Because it was grassy and wanted wear;',
        'Though as for that the passing there',
        'Had worn them really about the same,',
    ]

    stanza3_lines = [
        'And both that morning equally lay',
        'In leaves no step had trodden black.',
        'Oh, I kept the first for another day!',
        'Yet knowing how way leads on to way,',
        'I doubted if I should ever come back.',
    ]

    stanza4_lines = [
        'I shall be telling this with a sigh',
        'Somewhere ages and ages hence:',
        'Two roads diverged in a wood, and I\u2014',
        'I took the one less traveled by,',
        'And that has made all the difference.',
    ]

    stanzas = [stanza1_lines, stanza2_lines, stanza3_lines, stanza4_lines]

    for stanza_lines in stanzas:
        para = doc.add_paragraph()
        for i, line in enumerate(stanza_lines):
            run = para.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if i < len(stanza_lines) - 1:
                # Add soft line break (Shift+Enter) between lines within a stanza
                run.add_break()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
