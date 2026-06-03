"""
Initial Setup: Two-column newsletter document with 'News Highlights' section
Task ID: writer_rd_018
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_018'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_section_columns(section, num_cols, spacing_cm=0.5, separator=False):
    """Set section columns using low-level XML."""
    sect_pr = section._sectPr
    # Remove existing cols element if present
    for cols_elem in sect_pr.findall(qn('w:cols')):
        sect_pr.remove(cols_elem)

    cols = sect_pr.makeelement(qn('w:cols'), {})
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(int(Cm(spacing_cm))))  # EMU value
    if separator:
        cols.set(qn('w:sep'), '1')
    else:
        cols.set(qn('w:sep'), '0')
    cols.set(qn('w:equalWidth'), '1')
    sect_pr.append(cols)


def create_initial():
    doc = Document()

    # Set page margins for a nice newsletter look
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title heading
    heading = doc.add_heading('News Highlights', level=1)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.paragraph_format.space_after = Pt(12)

    # 6 paragraphs of realistic newsletter content
    paragraphs_text = [
        (
            "The city council approved a landmark urban renewal initiative on Tuesday, "
            "allocating $4.2 million toward the revitalization of the downtown waterfront "
            "district. Mayor Elena Vasquez described the project as a once-in-a-generation "
            "opportunity to transform the riverfront into a vibrant community hub, featuring "
            "pedestrian walkways, public art installations, and green spaces."
        ),
        (
            "Local tech startup NovaBridge Solutions announced a strategic partnership with "
            "Meridian Healthcare Systems to develop an AI-powered patient triage platform. "
            "CEO Rajesh Patel confirmed that the collaboration will bring 85 new engineering "
            "positions to the region over the next eighteen months, with an initial investment "
            "of $12 million dedicated to research and development."
        ),
        (
            "The Westfield Community Library celebrated its centennial anniversary with a "
            "weekend festival that drew over 3,000 visitors. Head librarian Doris Kaplan "
            "curated a special exhibition showcasing rare first editions and archival "
            "photographs from the library's founding in 1926. Local author Simone Delacroix "
            "delivered the keynote address on the enduring importance of public literacy."
        ),
        (
            "Construction on the new Greenview Metro Extension is now sixty percent complete, "
            "according to the Regional Transit Authority. Project manager Tomás Herrera "
            "reported that the 4.8-mile expansion remains on schedule for a spring opening, "
            "with three new stations connecting residential neighborhoods in the eastern "
            "corridor to the central business district."
        ),
        (
            "The annual Harvest Arts Festival returned to Elmwood Park last Saturday, "
            "featuring works by 120 regional artists and craftspeople. Festival director "
            "Clara Okonkwo highlighted the introduction of a live mural painting competition "
            "and an expanded culinary village with dishes from 18 local restaurants. "
            "Attendance surpassed expectations, with an estimated 7,500 visitors over two days."
        ),
        (
            "Riverside High School's robotics team, the Circuit Breakers, secured first place "
            "at the National STEM Innovation Challenge held in Denver last weekend. Coach "
            "Angela Whitmore credited months of intensive preparation and collaboration, "
            "noting that the team's autonomous navigation module outperformed entries from "
            "42 other schools across the country. The team will advance to the international "
            "finals in Munich this July."
        ),
    ]

    for text in paragraphs_text:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Liberation Serif'
        run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Set section to 2 columns, 0.5 cm spacing, NO separator
    set_section_columns(section, num_cols=2, spacing_cm=0.5, separator=False)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
