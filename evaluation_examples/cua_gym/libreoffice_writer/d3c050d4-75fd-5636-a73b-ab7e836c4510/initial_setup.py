"""
Initial Setup: Compensation Restructuring Proposal - raw data document
Task ID: writer_hr_071
Domain: libreoffice_writer

Creates a Writer document with the title 'Compensation_Restructuring_Proposal'
and raw data in text form. No tables, no structured sections, no page breaks.
The agent must transform this into a proper multi-section proposal.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_071'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title - just plain text, not a structured heading
    title_para = doc.add_paragraph()
    run = title_para.add_run("Compensation_Restructuring_Proposal")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph("")

    # Raw executive summary data as plain text
    doc.add_paragraph(
        "Executive Summary Notes: After reviewing compensation data across "
        "the organization, we recommend a comprehensive restructuring of our "
        "salary grades to remain competitive in the current labor market. "
        "Key findings indicate that 62% of our positions fall below the 50th "
        "percentile of market rates. The proposed restructuring would affect "
        "approximately 340 employees across 8 salary grades with an estimated "
        "annual budget increase of $2.4M (8.3% over current spend)."
    )

    doc.add_paragraph(
        "Key Recommendations: 1) Implement a new 8-grade salary structure aligned "
        "with market data. 2) Prioritize adjustments for roles more than 15% below "
        "market midpoint. 3) Phase implementation over 18 months to manage budget "
        "impact. 4) Establish annual market review cadence."
    )

    doc.add_paragraph("")

    # Raw market analysis data as text
    doc.add_paragraph(
        "Market Analysis Data - Benchmark Position Comparison:"
    )
    doc.add_paragraph(
        "The following data was collected from 5 salary surveys: "
        "Mercer Total Compensation Survey, Radford Technology Survey, "
        "Willis Towers Watson General Industry Survey, Culpepper Compensation Survey, "
        "and PayScale Industry Report."
    )

    doc.add_paragraph("")

    # Raw benchmark data in text form (10 positions)
    benchmark_data = [
        ("Software Engineer II", "$95,000", "$102,500", "$98,000", "$100,200", "$97,800", "$98,700"),
        ("Senior Software Engineer", "$128,000", "$135,000", "$131,500", "$133,200", "$129,800", "$131,500"),
        ("Engineering Manager", "$155,000", "$162,000", "$158,500", "$160,100", "$156,200", "$158,360"),
        ("Data Analyst", "$72,000", "$78,500", "$75,200", "$76,800", "$73,400", "$75,180"),
        ("Product Manager", "$118,000", "$125,000", "$121,500", "$123,000", "$119,800", "$121,460"),
        ("UX Designer", "$88,000", "$94,500", "$91,200", "$92,800", "$89,600", "$91,220"),
        ("HR Business Partner", "$85,000", "$91,000", "$88,500", "$89,200", "$86,800", "$88,100"),
        ("Financial Analyst", "$78,000", "$84,000", "$81,200", "$82,500", "$79,600", "$81,060"),
        ("Marketing Manager", "$105,000", "$112,000", "$108,500", "$110,200", "$106,800", "$108,500"),
        ("Operations Director", "$142,000", "$150,000", "$146,500", "$148,200", "$144,000", "$146,140"),
    ]

    for pos, mercer, radford, wtw, culpepper, payscale, avg in benchmark_data:
        doc.add_paragraph(
            f"{pos}: Mercer {mercer}, Radford {radford}, WTW {wtw}, "
            f"Culpepper {culpepper}, PayScale {payscale}, Average {avg}"
        )

    doc.add_paragraph("")

    # Raw salary structure data
    doc.add_paragraph("Proposed Salary Structure Data:")

    grade_data = [
        ("Grade 1", "$42,000", "$52,500", "$63,000", "50%"),
        ("Grade 2", "$55,000", "$68,750", "$82,500", "50%"),
        ("Grade 3", "$70,000", "$87,500", "$105,000", "50%"),
        ("Grade 4", "$88,000", "$110,000", "$132,000", "50%"),
        ("Grade 5", "$108,000", "$135,000", "$162,000", "50%"),
        ("Grade 6", "$132,000", "$165,000", "$198,000", "50%"),
        ("Grade 7", "$160,000", "$200,000", "$240,000", "50%"),
        ("Grade 8", "$195,000", "$243,750", "$292,500", "50%"),
    ]

    for grade, min_sal, mid, max_sal, spread in grade_data:
        doc.add_paragraph(
            f"{grade}: Minimum {min_sal}, Midpoint {mid}, Maximum {max_sal}, Spread {spread}"
        )

    doc.add_paragraph("")

    # Raw budget impact data
    doc.add_paragraph("Budget Impact Analysis Data:")

    dept_data = [
        ("Engineering", "$8,450,000", "$9,180,000", "$730,000", "8.6%"),
        ("Product", "$3,200,000", "$3,460,000", "$260,000", "8.1%"),
        ("Design", "$2,100,000", "$2,280,000", "$180,000", "8.6%"),
        ("Data Science", "$2,850,000", "$3,090,000", "$240,000", "8.4%"),
        ("Human Resources", "$1,950,000", "$2,110,000", "$160,000", "8.2%"),
        ("Finance", "$2,400,000", "$2,600,000", "$200,000", "8.3%"),
        ("Marketing", "$2,750,000", "$2,990,000", "$240,000", "8.7%"),
        ("Operations", "$3,150,000", "$3,410,000", "$260,000", "8.3%"),
    ]

    for dept, current, proposed, delta, pct in dept_data:
        doc.add_paragraph(
            f"{dept}: Current {current}, Proposed {proposed}, Delta {delta}, Change {pct}"
        )

    doc.add_paragraph(
        "Total: Current $26,850,000, Proposed $29,120,000, Delta $2,270,000, Change 8.5%"
    )

    doc.add_paragraph("")

    # Raw implementation timeline data
    doc.add_paragraph("Implementation Timeline Data:")
    doc.add_paragraph(
        "Phase 1 (Months 1-3): Market data validation and final grade mapping. "
        "Activities include survey data analysis, job evaluation committee meetings, "
        "grade structure finalization."
    )
    doc.add_paragraph(
        "Phase 2 (Months 4-6): Communication strategy and manager training. "
        "Activities include developing communication materials, conducting manager "
        "briefings, preparing individual impact statements."
    )
    doc.add_paragraph(
        "Phase 3 (Months 7-12): Phased salary adjustments rollout. "
        "Activities include processing first wave adjustments for employees more than "
        "15% below midpoint, second wave for 10-15% below, third wave for remainder."
    )
    doc.add_paragraph(
        "Phase 4 (Months 13-18): Monitoring and refinement. "
        "Activities include quarterly compensation reviews, turnover analysis, "
        "budget reconciliation, annual market benchmarking setup."
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Cross-reference notes: The salary structure in the proposed grades should "
        "align with the market analysis benchmarks. Budget impact figures derive from "
        "mapping current employees to proposed grades. Timeline phases correspond to "
        "the priority levels identified in the executive summary recommendations."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
