"""
Reward Script: Multi-section compensation proposal document
Task ID: writer_hr_071
Domain: libreoffice_writer
Scoring:
  Component 1: Structured headings (Title + Heading styles)    — 0.15
  Component 2: Market analysis table (11 rows x 7 cols)        — 0.20
  Component 3: Salary structure table (9 rows x 5 cols)        — 0.20
  Component 4: Budget impact table (10 rows x 5 cols)          — 0.15
  Component 5: Gantt-style timeline table (5+ rows x 4+ cols)  — 0.15
  Component 6: Cross-references between sections               — 0.10
  Component 7: Chart placeholders present                      — 0.05
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_071'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            import time
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------------------------------------------------------------
    # Component 1: Structured headings — Title + Heading 1/2 (0.15)
    # Initial has ZERO headings (all Normal). Golden has Title + multiple Heading 1/2.
    # ---------------------------------------------------------------
    try:
        heading_styles = set()
        for p in doc.paragraphs:
            sn = p.style.name if p.style else ""
            if sn == "Title" or sn.startswith("Heading"):
                heading_styles.add(sn)

        has_title = "Title" in heading_styles
        has_h1 = any(s == "Heading 1" for s in heading_styles)
        has_h2 = any(s == "Heading 2" for s in heading_styles)

        # Count Heading 1 paragraphs (expect at least 5 sections)
        h1_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == "Heading 1")

        if has_title and has_h1 and has_h2 and h1_count >= 4:
            print(f"PASS: Component 1 — Structured headings: Title + {h1_count} H1 + H2 present (0.15 pts)")
            total_score += 0.15
        elif has_h1 and h1_count >= 3:
            print(f"PARTIAL: Component 1 — Has H1 headings ({h1_count}) but missing Title or H2 (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 1 — Expected structured headings, found styles: {heading_styles}, H1 count: {h1_count}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: Market analysis table — 10 benchmark positions x 5 surveys + avg (0.20)
    # Initial has 0 tables. Golden has table 0 = 11 rows x 7 cols.
    # ---------------------------------------------------------------
    try:
        tables = doc.tables
        if len(tables) < 1:
            print(f"FAIL: Component 2 — No tables found (expected market analysis table)")
        else:
            # Find a table that looks like the market analysis table
            # Header should contain 'Position' and survey names
            market_table = None
            for t in tables:
                if len(t.rows) >= 8 and len(t.columns) >= 6:
                    header_text = " ".join(c.text.strip().lower() for c in t.rows[0].cells)
                    if "position" in header_text or "mercer" in header_text or "survey" in header_text:
                        market_table = t
                        break

            if market_table is None:
                # Fall back: first table with >= 10 data rows and >= 6 columns
                for t in tables:
                    if len(t.rows) >= 10 and len(t.columns) >= 6:
                        market_table = t
                        break

            if market_table is None:
                print(f"FAIL: Component 2 — No market analysis table found (need >=10 rows, >=6 cols)")
            else:
                nrows = len(market_table.rows)
                ncols = len(market_table.columns)

                # Check dimensions (header + 10 positions = 11 rows, 7 cols)
                dim_ok = nrows >= 11 and ncols >= 6
                # Check that some expected positions are present
                position_texts = [market_table.cell(r, 0).text.strip().lower() for r in range(1, min(nrows, 12))]
                expected_positions = ["software engineer", "data analyst", "product manager", "ux designer", "financial analyst"]
                positions_found = sum(1 for exp in expected_positions if any(exp in pt for pt in position_texts))

                # Check that salary values appear (dollar amounts)
                has_salary_data = False
                for r in range(1, min(nrows, 3)):
                    for c in range(1, min(ncols, 4)):
                        val = market_table.cell(r, c).text.strip()
                        if "$" in val or val.replace(",", "").replace(".", "").isdigit():
                            has_salary_data = True
                            break
                    if has_salary_data:
                        break

                if dim_ok and positions_found >= 3 and has_salary_data:
                    print(f"PASS: Component 2 — Market analysis table {nrows}x{ncols}, {positions_found}/5 key positions found, salary data present (0.20 pts)")
                    total_score += 0.20
                elif dim_ok and positions_found >= 2:
                    print(f"PARTIAL: Component 2 — Table dims OK but only {positions_found} positions found (0.12 pts)")
                    total_score += 0.12
                else:
                    print(f"FAIL: Component 2 — Market table incomplete: dim_ok={dim_ok} ({nrows}x{ncols}), positions={positions_found}, salary_data={has_salary_data}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Salary structure table — 8 grades x 5 cols (0.20)
    # Columns: Grade, Min, Midpoint, Max, Spread
    # Initial has 0 tables.
    # ---------------------------------------------------------------
    try:
        salary_table = None
        for t in tables:
            if len(t.rows) >= 8 and len(t.columns) >= 4:
                header_text = " ".join(c.text.strip().lower() for c in t.rows[0].cells)
                if "grade" in header_text and ("min" in header_text or "midpoint" in header_text):
                    salary_table = t
                    break

        if salary_table is None:
            print(f"FAIL: Component 3 — No salary structure table found")
        else:
            nrows = len(salary_table.rows)
            ncols = len(salary_table.columns)

            # Check for 8 grade rows (header + 8 data = 9 rows)
            grade_count = 0
            for r in range(1, nrows):
                cell_text = salary_table.cell(r, 0).text.strip().lower()
                if "grade" in cell_text:
                    grade_count += 1

            # Check for proper columns (Grade, Min, Midpoint, Max, Spread)
            header_cells = [salary_table.cell(0, c).text.strip().lower() for c in range(min(ncols, 6))]
            has_min = any("min" in h for h in header_cells)
            has_mid = any("mid" in h for h in header_cells)
            has_max = any("max" in h for h in header_cells)
            has_spread = any("spread" in h for h in header_cells)

            if grade_count >= 7 and has_min and has_mid and has_max and has_spread and ncols >= 5:
                print(f"PASS: Component 3 — Salary structure table {nrows}x{ncols}, {grade_count} grades, all columns present (0.20 pts)")
                total_score += 0.20
            elif grade_count >= 5 and ncols >= 4:
                print(f"PARTIAL: Component 3 — Salary table has {grade_count} grades, {ncols} cols (0.12 pts)")
                total_score += 0.12
            else:
                print(f"FAIL: Component 3 — Salary table incomplete: grades={grade_count}, cols={ncols}, min={has_min}, mid={has_mid}, max={has_max}, spread={has_spread}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # ---------------------------------------------------------------
    # Component 4: Budget impact table — departments with current vs proposed costs (0.15)
    # Initial has 0 tables.
    # ---------------------------------------------------------------
    try:
        budget_table = None
        for t in tables:
            if len(t.rows) >= 6 and len(t.columns) >= 4:
                header_text = " ".join(c.text.strip().lower() for c in t.rows[0].cells)
                if "department" in header_text and ("current" in header_text or "proposed" in header_text or "cost" in header_text or "delta" in header_text):
                    budget_table = t
                    break

        if budget_table is None:
            print(f"FAIL: Component 4 — No budget impact table found")
        else:
            nrows = len(budget_table.rows)
            ncols = len(budget_table.columns)

            # Check for department rows
            dept_keywords = ["engineering", "product", "design", "marketing", "finance", "operations", "human resources", "data"]
            dept_count = 0
            for r in range(1, nrows):
                cell_text = budget_table.cell(r, 0).text.strip().lower()
                if any(kw in cell_text for kw in dept_keywords):
                    dept_count += 1

            # Check for a total row
            has_total = False
            for r in range(1, nrows):
                cell_text = budget_table.cell(r, 0).text.strip().lower()
                if "total" in cell_text:
                    has_total = True
                    break

            if dept_count >= 6 and has_total and ncols >= 4:
                print(f"PASS: Component 4 — Budget impact table {nrows}x{ncols}, {dept_count} departments + total (0.15 pts)")
                total_score += 0.15
            elif dept_count >= 4:
                print(f"PARTIAL: Component 4 — Budget table has {dept_count} departments (0.08 pts)")
                total_score += 0.08
            else:
                print(f"FAIL: Component 4 — Budget table incomplete: depts={dept_count}, total={has_total}, cols={ncols}")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # ---------------------------------------------------------------
    # Component 5: Gantt-style timeline table — phases with time periods (0.15)
    # Initial has 0 tables. Golden has table with phases and month columns with shading.
    # ---------------------------------------------------------------
    try:
        gantt_table = None
        for t in tables:
            if len(t.rows) >= 4 and len(t.columns) >= 4:
                header_text = " ".join(c.text.strip().lower() for c in t.rows[0].cells)
                if "phase" in header_text and "month" in header_text:
                    gantt_table = t
                    break

        if gantt_table is None:
            # Fallback: look for table with "phase" in first column cells
            for t in tables:
                if len(t.rows) >= 4:
                    phase_count = sum(1 for r in range(min(len(t.rows), 6))
                                      if "phase" in t.cell(r, 0).text.strip().lower())
                    if phase_count >= 3:
                        gantt_table = t
                        break

        if gantt_table is None:
            print(f"FAIL: Component 5 — No Gantt-style timeline table found")
        else:
            nrows = len(gantt_table.rows)
            ncols = len(gantt_table.columns)

            # Check for phase rows
            phase_count = 0
            for r in range(1, nrows):
                cell_text = gantt_table.cell(r, 0).text.strip().lower()
                if "phase" in cell_text:
                    phase_count += 1

            # Check for Gantt-style shading (colored cells indicating active periods)
            shaded_cells = 0
            ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            for ri in range(1, nrows):
                for ci in range(2, ncols):  # skip Phase and Activities columns
                    cell_el = gantt_table.cell(ri, ci)._element
                    shd = cell_el.find(f".//{{{ns_w}}}shd")
                    if shd is not None:
                        fill = shd.attrib.get(f"{{{ns_w}}}fill", "")
                        if fill and fill not in ("auto", "FFFFFF", "ffffff", "F2F2F2", "f2f2f2"):
                            shaded_cells += 1

            if phase_count >= 3 and ncols >= 5 and shaded_cells >= 3:
                print(f"PASS: Component 5 — Gantt timeline table {nrows}x{ncols}, {phase_count} phases, {shaded_cells} Gantt-shaded cells (0.15 pts)")
                total_score += 0.15
            elif phase_count >= 3 and ncols >= 4:
                print(f"PARTIAL: Component 5 — Timeline table found, {phase_count} phases, but limited Gantt shading ({shaded_cells}) (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 5 — Timeline table incomplete: phases={phase_count}, cols={ncols}, shading={shaded_cells}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # ---------------------------------------------------------------
    # Component 6: Cross-references between sections (0.10)
    # Initial has 0 cross-references. Golden has many "Section N" references.
    # ---------------------------------------------------------------
    try:
        cross_ref_count = 0
        sections_referenced = set()
        for p in doc.paragraphs:
            refs = re.findall(r"[Ss]ection\s+(\d)", p.text)
            cross_ref_count += len(refs)
            for r_num in refs:
                sections_referenced.add(int(r_num))

        if cross_ref_count >= 8 and len(sections_referenced) >= 4:
            print(f"PASS: Component 6 — {cross_ref_count} cross-references across {len(sections_referenced)} sections (0.10 pts)")
            total_score += 0.10
        elif cross_ref_count >= 3:
            print(f"PARTIAL: Component 6 — {cross_ref_count} cross-references, {len(sections_referenced)} unique sections (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — Only {cross_ref_count} cross-references found (expected >=8)")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # ---------------------------------------------------------------
    # Component 7: Chart placeholders present (0.05)
    # Initial has none. Golden has at least 2 "[Chart Placeholder: ...]" paragraphs.
    # ---------------------------------------------------------------
    try:
        chart_placeholders = 0
        for p in doc.paragraphs:
            if "[Chart Placeholder" in p.text or "[chart placeholder" in p.text.lower():
                chart_placeholders += 1

        if chart_placeholders >= 2:
            print(f"PASS: Component 7 — {chart_placeholders} chart placeholders found (0.05 pts)")
            total_score += 0.05
        elif chart_placeholders >= 1:
            print(f"PARTIAL: Component 7 — {chart_placeholders} chart placeholder found (0.03 pts)")
            total_score += 0.03
        else:
            print(f"FAIL: Component 7 — No chart placeholders found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
