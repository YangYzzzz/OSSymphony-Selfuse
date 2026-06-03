"""
Reward Script: Apply bullet list using U+2192 arrow as bullet character
Task ID: writer_list_052
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All 5 paragraphs have numbering (numPr) applied at level 0
  Component 2 (0.3 pts): The numbering definition referenced uses U+2192 (→) as bullet char with numFmt=bullet
  Component 3 (0.2 pts): All 5 original text strings are preserved in the document
"""

import os
import zipfile
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_list_052'
FILE_PATH = f'{WORKDIR}/Desktop/navigation.docx'

EXPECTED_TEXTS = [
    "Click on Settings in the top menu bar",
    "Select Account Preferences from the dropdown",
    "Navigate to the Security section",
    "Enable two-factor authentication toggle",
    "Save changes and confirm with your password",
]

ARROW_CHAR = '\u2192'  # U+2192 rightwards arrow


def get_num_id_to_abstract_map(numbering_xml: str) -> dict:
    """Parse numbering.xml and return {numId: abstractNumId} mapping."""
    mapping = {}
    for m in re.finditer(r'<w:num\s+w:numId="(\d+)"[^>]*>.*?<w:abstractNumId\s+w:val="(\d+)"', numbering_xml, re.DOTALL):
        num_id = m.group(1)
        abstract_id = m.group(2)
        mapping[num_id] = abstract_id
    return mapping


def get_abstract_num_lvl_text(numbering_xml: str, abstract_num_id: str, ilvl: str = '0') -> str:
    """Extract lvlText value for a given abstractNumId and ilvl."""
    # Match the abstractNum block for the given ID
    pattern = r'<w:abstractNum\s+w:abstractNumId="' + re.escape(abstract_num_id) + r'".*?</w:abstractNum>'
    m = re.search(pattern, numbering_xml, re.DOTALL)
    if not m:
        return ''
    abstract_block = m.group(0)
    # Within the abstractNum block, find the level matching ilvl
    lvl_pattern = r'<w:lvl\s+w:ilvl="' + re.escape(ilvl) + r'".*?</w:lvl>'
    m2 = re.search(lvl_pattern, abstract_block, re.DOTALL)
    if not m2:
        return ''
    lvl_block = m2.group(0)
    # Extract lvlText value
    m3 = re.search(r'<w:lvlText\s+w:val="([^"]*)"', lvl_block)
    if not m3:
        return ''
    return m3.group(1)


def get_abstract_num_fmt(numbering_xml: str, abstract_num_id: str, ilvl: str = '0') -> str:
    """Extract numFmt value for a given abstractNumId and ilvl."""
    pattern = r'<w:abstractNum\s+w:abstractNumId="' + re.escape(abstract_num_id) + r'".*?</w:abstractNum>'
    m = re.search(pattern, numbering_xml, re.DOTALL)
    if not m:
        return ''
    abstract_block = m.group(0)
    lvl_pattern = r'<w:lvl\s+w:ilvl="' + re.escape(ilvl) + r'".*?</w:lvl>'
    m2 = re.search(lvl_pattern, abstract_block, re.DOTALL)
    if not m2:
        return ''
    lvl_block = m2.group(0)
    m3 = re.search(r'<w:numFmt\s+w:val="([^"]*)"', lvl_block)
    if not m3:
        return ''
    return m3.group(1)


def verify_task(file_path: str) -> float:
    """
    Verify that all 5 navigation paragraphs have been formatted as a bullet list
    using U+2192 (rightwards arrow) as the bullet character.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Pre-condition: file must exist and be a valid docx
    if not os.path.exists(file_path):
        print(f"CRITICAL: File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load docx {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Read numbering.xml for bullet character checks
    numbering_xml = ''
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            if 'word/numbering.xml' in z.namelist():
                numbering_xml = z.read('word/numbering.xml').decode('utf-8')
    except Exception as e:
        print(f"WARN: Could not read numbering.xml: {e}")

    # --- Component 1: All 5 paragraphs have numPr (numbering applied) at ilvl=0 (0.5 pts) ---
    # This must FAIL on initial_env (no numPr present) and PASS on golden_env
    try:
        paragraphs_with_numbering = 0
        paragraphs_at_level0 = 0
        para_num_ids = set()

        for para in doc.paragraphs:
            pPr = para._p.find(qn('w:pPr'))
            if pPr is None:
                continue
            numPr = pPr.find(qn('w:numPr'))
            if numPr is None:
                continue

            # Check text is one of our expected texts (to target the right paragraphs)
            para_text = para.text.strip()
            if para_text not in EXPECTED_TEXTS:
                continue

            ilvl_el = numPr.find(qn('w:ilvl'))
            numId_el = numPr.find(qn('w:numId'))

            if ilvl_el is not None and numId_el is not None:
                ilvl_val = ilvl_el.attrib.get(qn('w:val'), '')
                numId_val = numId_el.attrib.get(qn('w:val'), '')
                paragraphs_with_numbering += 1
                if ilvl_val == '0':
                    paragraphs_at_level0 += 1
                if numId_val:
                    para_num_ids.add(numId_val)

        if paragraphs_with_numbering == 5 and paragraphs_at_level0 == 5:
            print(f"PASS: Component 1 — All 5 paragraphs have numbering (numPr) applied at level 0 (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Expected 5 paragraphs with numPr at ilvl=0, "
                  f"found {paragraphs_with_numbering} with numPr, {paragraphs_at_level0} at ilvl=0")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # --- Component 2: The numbering uses U+2192 (→) as the bullet character with numFmt=bullet (0.3 pts) ---
    # This must FAIL on initial_env (no numbering definition for arrow bullet) and PASS on golden_env
    try:
        if not numbering_xml:
            print("FAIL: Component 2 — No numbering.xml found in docx")
        else:
            # Find which numId is used by the task paragraphs
            num_id_map = get_num_id_to_abstract_map(numbering_xml)

            # Check all numIds used by the paragraphs
            arrow_bullet_found = False
            checked_abstract_ids = set()

            if para_num_ids:
                for num_id in para_num_ids:
                    abstract_id = num_id_map.get(num_id, '')
                    if abstract_id in checked_abstract_ids:
                        continue
                    checked_abstract_ids.add(abstract_id)

                    lvl_text = get_abstract_num_lvl_text(numbering_xml, abstract_id, '0')
                    num_fmt = get_abstract_num_fmt(numbering_xml, abstract_id, '0')

                    print(f"  DEBUG: numId={num_id} -> abstractNumId={abstract_id}, "
                          f"lvlText={repr(lvl_text)}, numFmt={repr(num_fmt)}")

                    if ARROW_CHAR in lvl_text and num_fmt == 'bullet':
                        arrow_bullet_found = (ARROW_CHAR in lvl_text and num_fmt == 'bullet')
                        break

            if arrow_bullet_found:
                print(f"PASS: Component 2 — Numbering uses U+2192 (→) as bullet character with numFmt=bullet (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — No numbering definition found with U+2192 (→) as bullet character. "
                      f"para_num_ids={para_num_ids}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # --- Component 3: All 5 original text strings are preserved (0.2 pts) ---
    # This checks text integrity — the task should not alter text content
    # NOTE: This is structured as part of a task-change compound check:
    # the bullet formatting is applied AND the original text is preserved.
    # On initial_env this will PASS (text is there but no bullets), but
    # this component is only awarded together with component 1 gating logic.
    # We implement it as a standalone check but ensure it doesn't award points
    # for pre-existing conditions by requiring the list to also be applied (para_num_ids set).
    try:
        if paragraphs_at_level0 == 0:
            # No bullet formatting was applied — skip awarding these points
            print("SKIP: Component 3 — Text preservation not scored because no bullet formatting was applied")
        else:
            doc_texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            all_present = all(expected in doc_texts for expected in EXPECTED_TEXTS)

            if all_present:
                print(f"PASS: Component 3 — All 5 original navigation text strings are preserved (0.2 pts)")
                total_score += 0.2
            else:
                missing = [t for t in EXPECTED_TEXTS if t not in doc_texts]
                print(f"FAIL: Component 3 — Missing text strings: {missing}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
