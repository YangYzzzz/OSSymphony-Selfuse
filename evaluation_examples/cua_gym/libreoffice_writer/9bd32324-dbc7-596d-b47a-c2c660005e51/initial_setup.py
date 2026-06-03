"""
Initial Setup: Mail merge document with two data sources
Task ID: writer_mt_049
Domain: libreoffice_writer

Creates a Policy_Update.docx with:
- A letterhead area with static company info placeholders
- A recipient address block with static placeholders
- Body content about a policy update
- Two CSV data source files: Employees.csv and CompanyInfo.csv
The document does NOT contain any merge fields yet (that's the task).
"""

import os
import shlex
import subprocess
import time
import csv
import random

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_049'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

def create_employees_csv():
    """Create Employees data source with 100 records."""
    filepath = f'{WORKDIR}/Employees.csv'
    headers = ['EmpName', 'EmpAddress', 'EmpCity', 'EmpState', 'EmpZip']

    first_names = [
        'Sarah', 'Marcus', 'Elena', 'David', 'Priya', 'James', 'Mei', 'Carlos',
        'Aisha', 'Robert', 'Yuki', 'Thomas', 'Fatima', 'William', 'Sonia',
        'Michael', 'Laura', 'Ahmed', 'Rachel', 'Kevin', 'Nina', 'Patrick',
        'Diana', 'Gregory', 'Alina', 'Christopher', 'Maria', 'Jonathan', 'Olivia', 'Daniel',
        'Samantha', 'Victor', 'Hannah', 'Felix', 'Isabella', 'Nathan', 'Sophia', 'Brandon',
        'Amelia', 'Trevor', 'Chloe', 'Derek', 'Victoria', 'Leonard', 'Grace',
        'Russell', 'Emma', 'Philip', 'Megan', 'Stuart'
    ]
    last_names = [
        'Chen', 'Johnson', 'Rodriguez', 'Kim', 'Patel', 'Williams', 'Tanaka', 'Garcia',
        'Hassan', 'Thompson', 'Nakamura', 'Anderson', 'Al-Rashid', 'Davis', 'Fernandez',
        'Wilson', 'Martinez', 'Singh', 'Miller', 'O\'Brien', 'Petrov', 'Moore',
        'Kowalski', 'Taylor', 'Volkov', 'Brown', 'Lopez', 'White', 'Zhang', 'Harris',
        'Jackson', 'Lee', 'Clark', 'Mueller', 'Rossi', 'Scott', 'Johansson', 'Adams',
        'Nguyen', 'Baker', 'Dubois', 'Morgan', 'Santos', 'Reed', 'Hoffman',
        'Cooper', 'Yamamoto', 'Bell', 'Costa', 'Hughes'
    ]
    streets = [
        'Oak Street', 'Maple Avenue', 'Pine Road', 'Cedar Lane', 'Elm Boulevard',
        'Birch Drive', 'Walnut Court', 'Spruce Way', 'Ash Terrace', 'Willow Place',
        'Cherry Hill Road', 'Magnolia Drive', 'Sycamore Lane', 'Chestnut Street', 'Poplar Avenue'
    ]
    cities = [
        ('Austin', 'TX'), ('Seattle', 'WA'), ('Denver', 'CO'), ('Portland', 'OR'),
        ('Chicago', 'IL'), ('Boston', 'MA'), ('Atlanta', 'GA'), ('Phoenix', 'AZ'),
        ('Nashville', 'TN'), ('Minneapolis', 'MN'), ('Charlotte', 'NC'), ('San Diego', 'CA'),
        ('Columbus', 'OH'), ('Indianapolis', 'IN'), ('Tampa', 'FL')
    ]

    random.seed(42)
    rows = []
    for i in range(100):
        fn = first_names[i % len(first_names)]
        ln = last_names[i % len(last_names)]
        name = f'{fn} {ln}'
        addr = f'{random.randint(100, 9999)} {random.choice(streets)}'
        city, state = random.choice(cities)
        zipcode = f'{random.randint(10000, 99999)}'
        rows.append([name, addr, city, state, zipcode])

    with open(filepath, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(rows)
    print(f'Created: {filepath} ({len(rows)} records)')

def create_company_csv():
    """Create CompanyInfo data source with 1 record."""
    filepath = f'{WORKDIR}/CompanyInfo.csv'
    headers = ['CompanyName', 'HQAddress', 'HQCity', 'Phone', 'CEO']
    data = [
        'Nextera Global Solutions',
        '4500 Innovation Parkway, Suite 200',
        'San Francisco, CA 94105',
        '(415) 555-8200',
        'Margaret Thornton'
    ]
    with open(filepath, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerow(data)
    print(f'Created: {filepath}')

def create_document():
    """Create the initial Policy_Update.docx without merge fields."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Letterhead Area (static company info - to be replaced with merge fields) ---
    p_company = doc.add_paragraph()
    p_company.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_company.paragraph_format.space_after = Pt(2)
    run = p_company.add_run('Nextera Global Solutions')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p_addr = doc.add_paragraph()
    p_addr.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_addr.paragraph_format.space_after = Pt(2)
    run = p_addr.add_run('4500 Innovation Parkway, Suite 200')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_city = doc.add_paragraph()
    p_city.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_city.paragraph_format.space_after = Pt(2)
    run = p_city.add_run('San Francisco, CA 94105')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_phone = doc.add_paragraph()
    p_phone.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_phone.paragraph_format.space_after = Pt(12)
    run = p_phone.add_run('Phone: (415) 555-8200')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Horizontal line ---
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(12)
    run = p_line.add_run('_' * 75)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # --- Date ---
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_after = Pt(12)
    run = p_date.add_run('March 15, 2026')
    run.font.size = Pt(11)

    # --- Recipient Address Block (static - to be replaced with merge fields) ---
    p_recip_name = doc.add_paragraph()
    p_recip_name.paragraph_format.space_after = Pt(0)
    run = p_recip_name.add_run('[Employee Name]')
    run.font.size = Pt(11)

    p_recip_addr = doc.add_paragraph()
    p_recip_addr.paragraph_format.space_after = Pt(0)
    run = p_recip_addr.add_run('[Employee Address]')
    run.font.size = Pt(11)

    p_recip_city = doc.add_paragraph()
    p_recip_city.paragraph_format.space_after = Pt(12)
    run = p_recip_city.add_run('[City, State ZIP]')
    run.font.size = Pt(11)

    # --- Greeting ---
    p_greeting = doc.add_paragraph()
    p_greeting.paragraph_format.space_after = Pt(6)
    run = p_greeting.add_run('Dear Employee,')
    run.font.size = Pt(11)

    # --- Body paragraphs ---
    body_text = [
        'We are writing to inform you of important updates to the company\'s Remote Work and Flexible Schedule Policy, effective April 1, 2026. These changes reflect our commitment to supporting a healthy work-life balance while maintaining operational excellence across all departments.',

        'The updated policy introduces a hybrid work model that allows eligible employees to work remotely up to three days per week, subject to manager approval and departmental requirements. Employees must maintain core availability hours between 10:00 AM and 3:00 PM in their local time zone to ensure seamless collaboration across teams.',

        'In addition, we are expanding our wellness benefits program to include a $500 annual stipend for home office equipment and ergonomic furniture. This stipend can be used for items such as standing desks, ergonomic chairs, monitors, and noise-canceling headphones. Receipts must be submitted through the employee expense portal within 30 days of purchase.',

        'For employees who prefer to work primarily from the office, we have redesigned our workspace to include dedicated quiet zones, collaborative meeting spaces, and refreshed break areas with complimentary coffee and snacks. Building access cards will be reactivated for all employees by March 25, 2026.',

        'Please review the full policy document on the company intranet under Human Resources > Policies > Remote Work Policy 2026. If you have any questions or need clarification, please contact your HR Business Partner or email hr-support@nexteraglobal.com.',

        'We appreciate your continued dedication and look forward to supporting your success in this new work environment.'
    ]

    for text in body_text:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # --- Closing ---
    p_closing = doc.add_paragraph()
    p_closing.paragraph_format.space_before = Pt(12)
    p_closing.paragraph_format.space_after = Pt(4)
    run = p_closing.add_run('Sincerely,')
    run.font.size = Pt(11)

    p_sig_name = doc.add_paragraph()
    p_sig_name.paragraph_format.space_before = Pt(24)
    p_sig_name.paragraph_format.space_after = Pt(0)
    run = p_sig_name.add_run('Margaret Thornton')
    run.bold = True
    run.font.size = Pt(11)

    p_sig_title = doc.add_paragraph()
    p_sig_title.paragraph_format.space_after = Pt(0)
    run = p_sig_title.add_run('Chief Executive Officer')
    run.font.size = Pt(11)

    p_sig_company = doc.add_paragraph()
    p_sig_company.paragraph_format.space_after = Pt(0)
    run = p_sig_company.add_run('Nextera Global Solutions')
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Created document: {OUTPUT}')

# Execute
create_employees_csv()
create_company_csv()
create_document()

# Launch LibreOffice Writer with the document
launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')
