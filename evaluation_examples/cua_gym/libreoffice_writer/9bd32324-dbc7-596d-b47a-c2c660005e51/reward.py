"""
Reward Script: Mail merge with two data sources (Employees + CompanyInfo)
Task ID: writer_mt_049
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.30): CompanyInfo merge fields in letterhead area
  - Component 2 (0.30): Employee merge fields in recipient address area
  - Component 3 (0.20): Both data sources coexist in same document
  - Component 4 (0.20): Salutation uses EmpName + correct total field count
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_049'

# Namespace for Word XML
WNS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# Expected merge field names from each data source
COMPANY_FIELDS = {'CompanyName', 'HQAddress', 'HQCity', 'Phone'}
EMPLOYEE_FIELDS = {'EmpName', 'EmpAddress', 'EmpCity', 'EmpState', 'EmpZip'}


def extract_merge_fields(doc):
    """Extract all MERGEFIELD names from the document, returning a list of (para_index, field_name) tuples."""
    fields = []
    for i, para in enumerate(doc.paragraphs):
        instr_texts = para._element.findall('.//w:instrText', WNS)
        for it in instr_texts:
            text = it.text.strip() if it.text else ''
            if text.startswith('MERGEFIELD'):
                # Extract field name: "MERGEFIELD FieldName" or "MERGEFIELD FieldName \\* MERGEFORMAT"
                parts = text.split()
                if len(parts) >= 2:
                    fields.append((i, parts[1]))
    return fields


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract all merge fields
    all_fields = extract_merge_fields(doc)
    field_names = [name for _, name in all_fields]
    unique_field_names = set(field_names)

    print(f"INFO: Found {len(all_fields)} merge field instances, unique names: {unique_field_names}")

    # Component 1: CompanyInfo merge fields in letterhead (0.30 points)
    # The letterhead is in paragraphs 0-3 (before the separator line).
    # Expected: CompanyName, HQAddress, HQCity, Phone
    try:
        letterhead_fields = set()
        for para_idx, field_name in all_fields:
            if para_idx <= 4 and field_name in COMPANY_FIELDS:
                letterhead_fields.add(field_name)

        found_count = len(letterhead_fields)
        expected_count = len(COMPANY_FIELDS)

        if found_count == expected_count:
            print(f"PASS: Component 1 -- All {expected_count} CompanyInfo fields in letterhead: {letterhead_fields} (0.30 pts)")
            total_score += 0.30
        elif found_count > 0:
            partial = round(0.30 * (found_count / expected_count), 2)
            print(f"PARTIAL: Component 1 -- {found_count}/{expected_count} CompanyInfo fields in letterhead: {letterhead_fields} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 -- No CompanyInfo merge fields found in letterhead area (paras 0-4)")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Employee merge fields in recipient address (0.30 points)
    # The recipient address is in paragraphs 6-8 (after date, before salutation).
    # Expected: EmpName, EmpAddress, EmpCity, EmpState, EmpZip
    try:
        recipient_fields = set()
        for para_idx, field_name in all_fields:
            if 5 <= para_idx <= 9 and field_name in EMPLOYEE_FIELDS:
                recipient_fields.add(field_name)

        found_count = len(recipient_fields)
        expected_count = len(EMPLOYEE_FIELDS)

        if found_count == expected_count:
            print(f"PASS: Component 2 -- All {expected_count} Employee fields in recipient area: {recipient_fields} (0.30 pts)")
            total_score += 0.30
        elif found_count > 0:
            partial = round(0.30 * (found_count / expected_count), 2)
            print(f"PARTIAL: Component 2 -- {found_count}/{expected_count} Employee fields in recipient area: {recipient_fields} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- No Employee merge fields found in recipient area (paras 5-9)")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Both data sources coexist in same document (0.20 points)
    # At least one CompanyInfo field AND at least one Employee field must be present
    try:
        has_company = len(unique_field_names & COMPANY_FIELDS) > 0
        has_employee = len(unique_field_names & EMPLOYEE_FIELDS) > 0

        if has_company and has_employee:
            print(f"PASS: Component 3 -- Both data sources referenced: CompanyInfo fields={unique_field_names & COMPANY_FIELDS}, Employee fields={unique_field_names & EMPLOYEE_FIELDS} (0.20 pts)")
            total_score += 0.20
        elif has_company or has_employee:
            print(f"PARTIAL: Component 3 -- Only one data source found. CompanyInfo={has_company}, Employee={has_employee} (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 -- No merge fields from either data source found")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Salutation uses EmpName + correct total field count (0.20 points)
    # The salutation (para ~9) should use EmpName, and total unique merge fields should be >= 9
    # (CompanyName, HQAddress, HQCity, Phone, EmpName, EmpAddress, EmpCity, EmpState, EmpZip)
    try:
        # Check salutation has EmpName
        salutation_has_empname = any(
            field_name == 'EmpName' and ('Dear' in doc.paragraphs[para_idx].text or 'dear' in doc.paragraphs[para_idx].text)
            for para_idx, field_name in all_fields
        )

        # Check total unique field coverage
        all_expected = COMPANY_FIELDS | EMPLOYEE_FIELDS
        coverage = len(unique_field_names & all_expected)
        full_coverage = coverage == len(all_expected)

        if salutation_has_empname and full_coverage:
            print(f"PASS: Component 4 -- Salutation uses EmpName and all {len(all_expected)} expected fields present (0.20 pts)")
            total_score += 0.20
        elif salutation_has_empname:
            print(f"PARTIAL: Component 4 -- Salutation uses EmpName but only {coverage}/{len(all_expected)} fields present (0.10 pts)")
            total_score += 0.10
        elif full_coverage:
            print(f"PARTIAL: Component 4 -- All fields present but salutation missing EmpName (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 -- Salutation missing EmpName and only {coverage}/{len(all_expected)} fields present")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
