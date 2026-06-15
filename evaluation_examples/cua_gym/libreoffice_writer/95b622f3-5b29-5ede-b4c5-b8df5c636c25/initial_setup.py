"""
Initial Setup: Comprehensive Guide with Heading 1, 2, and 3 paragraphs (no TOC)
Task ID: osworld_writer_toc_generation_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_toc_generation_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Introduction Section (Heading 1) ---
    doc.add_heading('Introduction to Cloud Computing', level=1)
    doc.add_paragraph(
        'Cloud computing has revolutionized the way organizations manage and deploy their IT infrastructure. '
        'By leveraging scalable resources on demand, businesses can reduce capital expenditure while improving '
        'flexibility and resilience. This guide provides a comprehensive overview of cloud computing concepts, '
        'architectures, and best practices.'
    )

    # --- Overview (Heading 2) ---
    doc.add_heading('What is Cloud Computing?', level=2)
    doc.add_paragraph(
        'Cloud computing refers to the delivery of computing services—including servers, storage, databases, '
        'networking, software, analytics, and intelligence—over the Internet to offer faster innovation, '
        'flexible resources, and economies of scale. You typically pay only for cloud services you use, '
        'helping you lower operating costs and run your infrastructure more efficiently.'
    )

    doc.add_heading('History and Evolution', level=2)
    doc.add_paragraph(
        'The concept of cloud computing dates back to the 1960s when John McCarthy proposed that '
        'computing may someday be organized as a public utility. Over the following decades, the development '
        'of virtualization technology, the proliferation of the Internet, and increasing bandwidth '
        'capabilities laid the groundwork for modern cloud services.'
    )

    # --- Sub-sections (Heading 3) ---
    doc.add_heading('Early Mainframe Era', level=3)
    doc.add_paragraph(
        'During the 1960s and 1970s, large organizations relied on mainframe computers. '
        'Time-sharing systems allowed multiple users to access a single machine simultaneously, '
        'which was an early precursor to multi-tenant cloud computing.'
    )

    doc.add_heading('Virtualization Breakthrough', level=3)
    doc.add_paragraph(
        'VMware and other companies pioneered hardware virtualization in the late 1990s, '
        'enabling a single physical server to run multiple virtual machines. '
        'This innovation dramatically improved resource utilization and laid the foundation '
        'for Infrastructure as a Service (IaaS).'
    )

    # --- Service Models (Heading 1) ---
    doc.add_heading('Cloud Service Models', level=1)
    doc.add_paragraph(
        'Cloud services are broadly categorized into three primary service models. Understanding '
        'each model helps organizations select the appropriate level of control, flexibility, '
        'and management for their specific needs.'
    )

    doc.add_heading('Infrastructure as a Service (IaaS)', level=2)
    doc.add_paragraph(
        'IaaS provides virtualized computing resources over the internet. Providers like Amazon Web Services, '
        'Microsoft Azure, and Google Cloud Platform offer virtual machines, storage, and networking on demand. '
        'Customers manage the operating systems, middleware, and applications, while the provider manages '
        'the underlying hardware and virtualization layer.'
    )

    doc.add_heading('Platform as a Service (PaaS)', level=2)
    doc.add_paragraph(
        'PaaS delivers hardware and software tools over the Internet, typically used for application '
        'development. A PaaS provider hosts the hardware and software on its own infrastructure, '
        'freeing developers from having to install in-house hardware and software for developing or '
        'running new applications.'
    )

    doc.add_heading('Software as a Service (SaaS)', level=2)
    doc.add_paragraph(
        'SaaS is a method for delivering software applications over the Internet on demand and typically '
        'on a subscription basis. With SaaS, cloud providers host and manage the software application '
        'and underlying infrastructure and handle any maintenance. Examples include Salesforce, '
        'Microsoft Office 365, and Google Workspace.'
    )

    # --- Heading 3 under Service Models ---
    doc.add_heading('Choosing the Right Model', level=3)
    doc.add_paragraph(
        'Organizations must assess their technical capabilities, compliance requirements, and '
        'budget constraints when selecting a service model. A hybrid approach using multiple '
        'models is common in enterprise environments.'
    )

    # --- Deployment Models (Heading 1) ---
    doc.add_heading('Cloud Deployment Models', level=1)
    doc.add_paragraph(
        'Beyond service models, cloud solutions can be deployed in several ways depending on '
        'the organization\'s requirements for control, security, and cost.'
    )

    doc.add_heading('Public Cloud', level=2)
    doc.add_paragraph(
        'Public cloud services are owned and operated by third-party cloud service providers, '
        'which deliver their computing resources like servers and storage over the Internet. '
        'Microsoft Azure, Amazon Web Services (AWS), and Google Cloud Platform are examples '
        'of public cloud providers. In a public cloud, all hardware, software, and other '
        'supporting infrastructure is owned and managed by the cloud provider.'
    )

    doc.add_heading('Private Cloud', level=2)
    doc.add_paragraph(
        'A private cloud refers to cloud computing resources used exclusively by a single '
        'business or organization. A private cloud can be physically located at the organization\'s '
        'on-site datacenter, or it can be hosted by a third-party service provider. In a private cloud, '
        'the services and infrastructure are always maintained on a private network and the hardware '
        'and software are dedicated solely to the organization.'
    )

    doc.add_heading('Hybrid Cloud', level=2)
    doc.add_paragraph(
        'Hybrid clouds combine public and private clouds, bound together by technology that allows '
        'data and applications to be shared between them. By allowing data and applications to move '
        'between private and public clouds, a hybrid cloud gives your business greater flexibility '
        'and more deployment options.'
    )

    # Heading 3 under Deployment Models ---
    doc.add_heading('Multi-Cloud Strategies', level=3)
    doc.add_paragraph(
        'Many enterprises adopt a multi-cloud strategy, leveraging services from multiple cloud '
        'providers to avoid vendor lock-in, improve resilience, and optimize costs. '
        'Managing multi-cloud environments requires robust governance, monitoring, and orchestration tools.'
    )

    doc.add_heading('Edge Computing Integration', level=3)
    doc.add_paragraph(
        'Edge computing extends the cloud by processing data closer to where it is generated. '
        'This approach reduces latency and bandwidth usage, making it ideal for IoT applications, '
        'autonomous vehicles, and real-time analytics.'
    )

    # --- Security (Heading 1) ---
    doc.add_heading('Cloud Security and Compliance', level=1)
    doc.add_paragraph(
        'Security remains one of the primary concerns for organizations considering cloud adoption. '
        'Understanding the shared responsibility model and implementing robust security practices '
        'are essential for protecting data and workloads in the cloud.'
    )

    doc.add_heading('Shared Responsibility Model', level=2)
    doc.add_paragraph(
        'The shared responsibility model defines the division of security responsibilities between '
        'the cloud provider and the customer. Providers secure the underlying infrastructure, '
        'while customers are responsible for securing their data, applications, and access management. '
        'Misunderstanding this boundary is a common source of cloud security incidents.'
    )

    doc.add_heading('Identity and Access Management', level=2)
    doc.add_paragraph(
        'IAM policies control who can access cloud resources and what actions they can perform. '
        'Implementing least-privilege access, multi-factor authentication, and regular access reviews '
        'are foundational IAM best practices that significantly reduce the attack surface.'
    )

    # Heading 3 under Security ---
    doc.add_heading('Encryption at Rest and in Transit', level=3)
    doc.add_paragraph(
        'All sensitive data should be encrypted both when stored and when transmitted. '
        'Cloud providers offer managed key management services (KMS) that integrate seamlessly '
        'with storage and database services, simplifying encryption implementation.'
    )

    doc.add_heading('Compliance Frameworks', level=3)
    doc.add_paragraph(
        'Organizations in regulated industries must align cloud deployments with relevant compliance '
        'frameworks such as GDPR, HIPAA, SOC 2, and ISO 27001. Most major cloud providers offer '
        'compliance certifications and tools to help customers achieve and maintain compliance.'
    )

    # --- Cost Management (Heading 1) ---
    doc.add_heading('Cost Management and Optimization', level=1)
    doc.add_paragraph(
        'Cloud cost management is an ongoing discipline that ensures organizations maximize the value '
        'of their cloud investment. Without proper governance, cloud costs can escalate rapidly '
        'due to unconstrained resource provisioning.'
    )

    doc.add_heading('Rightsizing Resources', level=2)
    doc.add_paragraph(
        'Rightsizing involves matching resource types and sizes to actual workload requirements. '
        'Organizations can significantly reduce costs by identifying and resizing overprovisioned '
        'instances, using cloud provider cost analysis tools and third-party monitoring solutions.'
    )

    doc.add_heading('Reserved Instances and Savings Plans', level=2)
    doc.add_paragraph(
        'Cloud providers offer significant discounts for committing to use a specific amount of '
        'resources over one or three-year terms. Reserved Instances on AWS and Azure Reserved VM Instances '
        'can provide savings of up to 72% compared to on-demand pricing, making them attractive for '
        'stable, predictable workloads.'
    )

    # Heading 3 under Cost Management ---
    doc.add_heading('Tagging and Allocation', level=3)
    doc.add_paragraph(
        'Consistent resource tagging enables accurate cost attribution across departments, '
        'projects, and environments. Establishing a tagging taxonomy and enforcing it via '
        'policy-as-code tools ensures that cost data is actionable and traceable.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
