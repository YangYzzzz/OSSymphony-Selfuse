"""
Reward Script: Insert a Table of Contents showing only Heading 1 and 2 levels with dotted leader lines
Task ID: osworld_writer_toc_generation_003
Domain: libreoffice_writer

Scoring:
  Component 1: TOC heading paragraph is present at the start of the document (0.25 pts)
  Component 2: TOC field instruction specifies levels 1-2 only (\o "1-2") (0.40 pts)
  Component 3: TOC styles (toc 1 / toc 2) have dotted leader lines (0.35 pts)
  Total: 1.0
"""

import os
import re

# python-docx is available in the VM environment
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_toc_generation_003'


def persist_app_state():
    """
    Send Ctrl+S to save any unsaved LibreOffice Writer edits before verifying.
    Best-effort: log warning on failure and continue.
    """
    try:
        import time
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the document contains a TOC with:
    1. A TOC Heading paragraph at the top of the document.
    2. A TOC field instruction that covers only levels 1-2 (no Heading 3).
    3. Dotted leader lines in the toc 1 and/or toc 2 styles.

    Returns a float between 0.0 and 1.0.
    """
    total_score = 0.0

    # Load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: TOC Heading paragraph exists at the start of the document
    # The task requires a TOC to be inserted — the TOC Heading paragraph
    # ('Table of Contents') must appear as the FIRST paragraph in the document.
    # Initial doc starts with 'Heading 1', so this fails on initial and passes
    # on golden.
    # (0.25 points)
    # -----------------------------------------------------------------------
    try:
        first_para = doc.paragraphs[0] if doc.paragraphs else None
        if first_para is not None and first_para.style.name == 'TOC Heading':
            print(f"PASS: Component 1 — First paragraph has 'TOC Heading' style "
                  f"(text={first_para.text[:60]!r}) (0.25 pts)")
            total_score += 0.25
        else:
            first_style = first_para.style.name if first_para else 'N/A'
            print(f"FAIL: Component 1 — Expected first paragraph style='TOC Heading', "
                  f"found style={first_style!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: TOC field instruction contains \o "1-2" (levels 1 and 2 only)
    # The task says "only Heading 1 and Heading 2 levels (not Heading 3)".
    # In Word/LibreOffice DOCX format, this is encoded as the TOC switch \o "1-2"
    # inside a w:instrText element containing the TOC field instruction.
    # The initial document has no TOC field instruction at all.
    # (0.40 points)
    # -----------------------------------------------------------------------
    try:
        # Find all w:instrText elements in the document body
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        body = doc.element.body
        instr_texts = body.findall('.//w:instrText', ns)

        toc_instr_found = False
        levels_correct = False
        toc_instr_text = None

        for instr in instr_texts:
            text = instr.text or ''
            if 'TOC' in text:
                toc_instr_found = True
                toc_instr_text = text.strip()
                # Check for \o "1-2" meaning levels 1 through 2 only
                # Pattern: \o followed by "1-2" (the quoted range)
                if re.search(r'\\o\s+"1-2"', text):
                    levels_correct = True
                break

        if not toc_instr_found:
            print("FAIL: Component 2 — No TOC field instruction found in document")
        elif not levels_correct:
            print(f"FAIL: Component 2 — TOC instruction found but not restricted to "
                  f"levels 1-2. Instruction: {toc_instr_text!r}")
        else:
            print(f"PASS: Component 2 — TOC instruction restricts to levels 1-2 "
                  f"(instruction: {toc_instr_text!r}) (0.40 pts)")
            total_score += 0.40

    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Dotted leader lines in TOC styles
    # The task requires "dotted leader lines between the heading title and
    # page number". In DOCX, this is encoded as a right-aligned tab stop with
    # w:leader="dot" in the paragraph's tab definitions (in the toc 1 / toc 2
    # style or directly in the paragraph's pPr).
    # We check: (a) style definitions for toc 1 / toc 2 have dot leader tabs,
    # OR (b) actual TOC entry paragraphs (with style toc 1 / toc 2) have dot
    # leader tab stops.
    # Initial doc does not define toc 1 / toc 2 styles with dot leaders.
    # (0.35 points)
    # -----------------------------------------------------------------------
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        dot_leader_found = False

        # Check style definitions first
        styles_part = doc.part.styles
        for style in styles_part:
            style_name = style.name.lower()
            if style_name in ('toc 1', 'toc 2'):
                # Look for tab stops with leader="dot" in style XML
                tabs = style._element.findall('.//w:tab', ns)
                for tab in tabs:
                    leader = tab.get(qn('w:leader'), '')
                    tab_val = tab.get(qn('w:val'), '')
                    if leader == 'dot' and tab_val in ('right', 'decimal'):
                        dot_leader_found = True
                        print(f"PASS: Component 3 — Style '{style.name}' has "
                              f"dotted leader tab stop (leader={leader!r}, "
                              f"val={tab_val!r}) (0.35 pts)")
                        break
                if dot_leader_found:
                    break

        # If not found in style definitions, check actual paragraph tab stops
        if not dot_leader_found:
            for para in doc.paragraphs:
                style_name = para.style.name.lower()
                if style_name in ('toc 1', 'toc 2'):
                    tabs = para._element.findall('.//w:tab', ns)
                    for tab in tabs:
                        leader = tab.get(qn('w:leader'), '')
                        tab_val = tab.get(qn('w:val'), '')
                        if leader == 'dot' and tab_val in ('right', 'decimal'):
                            dot_leader_found = True
                            print(f"PASS: Component 3 — Paragraph with style "
                                  f"'{para.style.name}' has dotted leader tab stop "
                                  f"(leader={leader!r}, val={tab_val!r}) (0.35 pts)")
                            break
                if dot_leader_found:
                    break

        if not dot_leader_found:
            print("FAIL: Component 3 — No dotted leader tab stops found in "
                  "toc 1 / toc 2 styles or TOC entry paragraphs")
        else:
            total_score += 0.35

    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Final score
    # -----------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path on VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
