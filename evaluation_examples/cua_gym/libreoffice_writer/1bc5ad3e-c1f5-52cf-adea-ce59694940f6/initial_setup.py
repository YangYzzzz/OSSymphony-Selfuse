"""
Initial Setup: Accept all tracked changes and enable track changes
Task ID: writer_rm_088
Domain: libreoffice_writer

Creates a manuscript document with 30 tracked changes from a first review round.
Track changes recording is currently OFF.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_088'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
}

AUTHOR = "Dr. Emily Torres"
DATE = "2025-11-18T14:30:00Z"

# Track change counter
_rev_id = [0]

def next_rev_id():
    _rev_id[0] += 1
    return str(_rev_id[0])


def make_rpr_xml(bold=False, italic=False, font_name=None, font_size_pt=None, color_rgb=None):
    """Build a <w:rPr> element string for run properties."""
    parts = []
    if font_name:
        parts.append(f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
    if font_size_pt:
        half_pt = int(font_size_pt * 2)
        parts.append(f'<w:sz w:val="{half_pt}"/><w:szCs w:val="{half_pt}"/>')
    if bold:
        parts.append('<w:b/><w:bCs/>')
    if italic:
        parts.append('<w:i/><w:iCs/>')
    if color_rgb:
        parts.append(f'<w:color w:val="{color_rgb}"/>')
    if not parts:
        return ''
    return '<w:rPr>' + ''.join(parts) + '</w:rPr>'


def xml_escape(text):
    """Escape XML special characters in text content."""
    return (text
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('"', '&quot;')
        .replace("'", '&apos;'))


def add_insertion(para_elem, text, author=AUTHOR, date=DATE, rpr_xml=''):
    """Add a tracked insertion (w:ins) to a paragraph element."""
    rid = next_rev_id()
    safe_text = xml_escape(text)
    ins_xml = (
        f'<w:ins {nsdecls("w")} w:id="{rid}" w:author="{author}" w:date="{date}">'
        f'  <w:r>{rpr_xml}<w:t xml:space="preserve">{safe_text}</w:t></w:r>'
        f'</w:ins>'
    )
    ins_elem = parse_xml(ins_xml)
    para_elem.append(ins_elem)
    return ins_elem


def add_deletion(para_elem, text, author=AUTHOR, date=DATE, rpr_xml=''):
    """Add a tracked deletion (w:del) to a paragraph element."""
    rid = next_rev_id()
    safe_text = xml_escape(text)
    del_xml = (
        f'<w:del {nsdecls("w")} w:id="{rid}" w:author="{author}" w:date="{date}">'
        f'  <w:r>{rpr_xml}<w:delText xml:space="preserve">{safe_text}</w:delText></w:r>'
        f'</w:del>'
    )
    del_elem = parse_xml(del_xml)
    para_elem.append(del_elem)
    return del_elem


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Title ---
    title = doc.add_heading('', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('The Silent Architecture of Urban Green Spaces:\nA Multidisciplinary Analysis')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # --- Author info ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Dr. Marcus Wei, Department of Urban Ecology\nRiverside Metropolitan University')
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()  # blank line

    # =====================================================
    # SECTION 1: ABSTRACT (with tracked changes)
    # =====================================================
    h1 = doc.add_heading('Abstract', level=1)

    # Paragraph 1 of abstract - has insertion (change 1) and deletion (change 2)
    p = doc.add_paragraph()
    run = p.add_run('Urban green spaces ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 1: insertion of "increasingly"
    add_insertion(p._element, 'increasingly ', author=AUTHOR, date=DATE)
    run2 = p.add_run('serve as critical infrastructure for ')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    # CHANGE 2: deletion of "the improvement of"
    add_deletion(p._element, 'the improvement of ', author=AUTHOR, date=DATE)
    # CHANGE 3: insertion of "improving"
    add_insertion(p._element, 'improving ', author=AUTHOR, date=DATE)
    run3 = p.add_run('public health outcomes in densely populated metropolitan areas. ')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    # CHANGE 4: insertion of new sentence
    add_insertion(p._element,
        'Recent longitudinal studies spanning 2018-2024 have reinforced these findings with compelling statistical evidence. ',
        author=AUTHOR, date=DATE)

    # Paragraph 2 of abstract
    p2 = doc.add_paragraph()
    run = p2.add_run('This paper examines the ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 5: deletion of "various"
    add_deletion(p2._element, 'various ', author=AUTHOR, date=DATE)
    # CHANGE 6: insertion of "multifaceted"
    add_insertion(p2._element, 'multifaceted ', author=AUTHOR, date=DATE)
    run2 = p2.add_run('relationships between urban vegetation density, air quality indices, and community wellness metrics across fourteen cities in North America and Europe.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # =====================================================
    # SECTION 2: INTRODUCTION
    # =====================================================
    doc.add_heading('1. Introduction', level=1)

    p3 = doc.add_paragraph()
    run = p3.add_run('The rapid expansion of urban areas throughout the ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 7: deletion of "twentieth"
    add_deletion(p3._element, 'twentieth ', author=AUTHOR, date=DATE)
    # CHANGE 8: insertion of "twenty-first"
    add_insertion(p3._element, 'twenty-first ', author=AUTHOR, date=DATE)
    run2 = p3.add_run('century has fundamentally altered the relationship between human populations and natural ecosystems. ')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run3 = p3.add_run('As of 2024, approximately 56% of the global population resides in urban settings, a figure projected to reach 68% by 2050 according to United Nations demographic forecasts.')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    p4 = doc.add_paragraph()
    run = p4.add_run('Green spaces within these urban environments, ranging from expansive metropolitan parks to ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 9: deletion of "small neighborhood gardens"
    add_deletion(p4._element, 'small neighborhood gardens', author=AUTHOR, date=DATE)
    # CHANGE 10: insertion of "intimate pocket gardens and rooftop installations"
    add_insertion(p4._element, 'intimate pocket gardens and rooftop installations', author=AUTHOR, date=DATE)
    run2 = p4.add_run(', provide essential ecosystem services that directly impact human wellbeing.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    # CHANGE 11: insertion of additional sentence
    add_insertion(p4._element,
        ' These services encompass air filtration, thermal regulation, stormwater management, and psychosocial restoration.',
        author=AUTHOR, date=DATE)

    p5 = doc.add_paragraph()
    run = p5.add_run('Despite ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 12: deletion of "growing"
    add_deletion(p5._element, 'growing ', author=AUTHOR, date=DATE)
    # CHANGE 13: insertion of "substantial and mounting"
    add_insertion(p5._element, 'substantial and mounting ', author=AUTHOR, date=DATE)
    run2 = p5.add_run('evidence supporting the health benefits of urban greenery, significant gaps persist in our understanding of the specific mechanisms through which vegetation density influences respiratory health, cardiovascular function, and mental wellness at the population level.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # =====================================================
    # SECTION 3: LITERATURE REVIEW
    # =====================================================
    doc.add_heading('2. Literature Review', level=1)

    p6 = doc.add_paragraph()
    run = p6.add_run('The foundational work of Ulrich (1984) established that visual exposure to natural settings ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 14: deletion of "can reduce"
    add_deletion(p6._element, 'can reduce ', author=AUTHOR, date=DATE)
    # CHANGE 15: insertion of "significantly reduces"
    add_insertion(p6._element, 'significantly reduces ', author=AUTHOR, date=DATE)
    run2 = p6.add_run('patient recovery times and stress biomarkers. ')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run3 = p6.add_run('Subsequent studies by Kaplan and Kaplan (1989) introduced Attention Restoration Theory, positing that natural environments facilitate cognitive recovery from directed attention fatigue.')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    p7 = doc.add_paragraph()
    # CHANGE 16: insertion of entire paragraph lead-in
    add_insertion(p7._element,
        'A groundbreaking meta-analysis by Thompson et al. (2022) synthesized data from 147 studies across 28 countries, ',
        author=AUTHOR, date=DATE)
    run = p7.add_run('confirming a robust inverse correlation between proximity to green spaces and incidence of cardiovascular disease, type 2 diabetes, and clinical depression.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p8 = doc.add_paragraph()
    run = p8.add_run('More recently, research has ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 17: deletion of "started to focus on"
    add_deletion(p8._element, 'started to focus on ', author=AUTHOR, date=DATE)
    # CHANGE 18: insertion of "pivoted toward"
    add_insertion(p8._element, 'pivoted toward ', author=AUTHOR, date=DATE)
    run2 = p8.add_run('quantifying the dose-response relationship between green space exposure and health outcomes, acknowledging that not all green spaces deliver equivalent benefits.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # =====================================================
    # SECTION 4: METHODOLOGY
    # =====================================================
    doc.add_heading('3. Methodology', level=1)

    p9 = doc.add_paragraph()
    run = p9.add_run('Our research employed a ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 19: deletion of "mixed-methods"
    add_deletion(p9._element, 'mixed-methods ', author=AUTHOR, date=DATE)
    # CHANGE 20: insertion of "convergent parallel mixed-methods"
    add_insertion(p9._element, 'convergent parallel mixed-methods ', author=AUTHOR, date=DATE)
    run2 = p9.add_run('design encompassing satellite imagery analysis, air quality monitoring, and longitudinal health surveys conducted between 2020 and 2024.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    p10 = doc.add_paragraph()
    run = p10.add_run('Vegetation density was measured using the Normalized Difference Vegetation Index (NDVI) derived from ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 21: deletion of "Landsat"
    add_deletion(p10._element, 'Landsat ', author=AUTHOR, date=DATE)
    # CHANGE 22: insertion of "Sentinel-2"
    add_insertion(p10._element, 'Sentinel-2 ', author=AUTHOR, date=DATE)
    run2 = p10.add_run('satellite data at 10-meter resolution. Air quality parameters, including PM2.5, PM10, NO2, and ground-level ozone concentrations, were obtained from a network of ')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    # CHANGE 23: deletion of "200"
    add_deletion(p10._element, '200 ', author=AUTHOR, date=DATE)
    # CHANGE 24: insertion of "342"
    add_insertion(p10._element, '342 ', author=AUTHOR, date=DATE)
    run3 = p10.add_run('monitoring stations across the study cities.')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    # =====================================================
    # SECTION 5: RESULTS
    # =====================================================
    doc.add_heading('4. Preliminary Results', level=1)

    p11 = doc.add_paragraph()
    run = p11.add_run('Initial analysis reveals a ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 25: deletion of "clear"
    add_deletion(p11._element, 'clear ', author=AUTHOR, date=DATE)
    # CHANGE 26: insertion of "statistically significant (p < 0.001)"
    add_insertion(p11._element, 'statistically significant (p < 0.001) ', author=AUTHOR, date=DATE)
    run2 = p11.add_run('negative correlation between NDVI values and PM2.5 concentrations across all fourteen study cities.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    p12 = doc.add_paragraph()
    run = p12.add_run('Neighborhoods with NDVI scores above 0.45 demonstrated ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 27: deletion of "lower rates"
    add_deletion(p12._element, 'lower rates ', author=AUTHOR, date=DATE)
    # CHANGE 28: insertion of "23% lower rates"
    add_insertion(p12._element, '23% lower rates ', author=AUTHOR, date=DATE)
    run2 = p12.add_run('of respiratory hospital admissions compared to areas with NDVI scores below 0.20, after controlling for socioeconomic variables, traffic density, and industrial proximity.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # =====================================================
    # SECTION 6: DISCUSSION
    # =====================================================
    doc.add_heading('5. Discussion', level=1)

    p13 = doc.add_paragraph()
    run = p13.add_run('These findings ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # CHANGE 29: deletion of "support"
    add_deletion(p13._element, 'support ', author=AUTHOR, date=DATE)
    # CHANGE 30: insertion of "substantially corroborate and extend"
    add_insertion(p13._element, 'substantially corroborate and extend ', author=AUTHOR, date=DATE)
    run2 = p13.add_run('prior research linking urban vegetation to improved respiratory outcomes. The magnitude of the observed effect sizes suggests that strategic urban greening initiatives could serve as cost-effective public health interventions, particularly in underserved communities with limited access to healthcare infrastructure.')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    p14 = doc.add_paragraph()
    run = p14.add_run('Further investigation is warranted to determine whether the threshold effects observed at NDVI 0.45 are consistent across different climate zones and urban morphologies. Additionally, the potential confounding influence of socioeconomic selection effects, whereby wealthier populations may self-select into greener neighborhoods, requires more rigorous instrumental variable analysis.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # =====================================================
    # REFERENCES (partial)
    # =====================================================
    doc.add_heading('References', level=1)

    refs = [
        'Kaplan, R., & Kaplan, S. (1989). The Experience of Nature: A Psychological Perspective. Cambridge University Press.',
        'Thompson, C. W., et al. (2022). Green space exposure and health outcomes: A systematic review and meta-analysis. Environmental Research Letters, 17(4), 043001.',
        'Ulrich, R. S. (1984). View through a window may influence recovery from surgery. Science, 224(4647), 420-421.',
        'United Nations. (2018). World Urbanization Prospects: The 2018 Revision. Department of Economic and Social Affairs.',
        'World Health Organization. (2016). Urban green spaces and health: A review of evidence. WHO Regional Office for Europe.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style='List Number')

    # --- Ensure track changes recording is OFF ---
    # We do NOT set w:trackChanges in document settings
    # The document has revision markup but recording is off

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Total tracked changes: {_rev_id[0]}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
