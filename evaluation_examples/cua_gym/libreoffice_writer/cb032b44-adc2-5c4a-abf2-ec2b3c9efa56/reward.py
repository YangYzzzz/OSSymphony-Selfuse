"""
Reward Script: Enable drop caps for first paragraph of each chapter
Task ID: writer_fs_016
Domain: libreoffice_writer
Scoring: 5 chapters x 0.2 pts each = 1.0. Each chapter's first body paragraph
         must have framePr with dropCap='drop' and lines='3'.
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_016'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Identify the first body paragraph after each Heading 1
    # These are the paragraphs that should have drop caps
    heading_indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == 'Heading 1':
            heading_indices.append(i)

    if len(heading_indices) == 0:
        print("FAIL: No Heading 1 paragraphs found in document")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found {len(heading_indices)} Heading 1 paragraphs at indices: {heading_indices}")

    # We expect 5 chapters; score each at 0.2
    expected_chapters = len(heading_indices)
    points_per_chapter = 1.0 / max(expected_chapters, 1)

    for chapter_num, h_idx in enumerate(heading_indices, 1):
        # The first body paragraph after this heading
        body_idx = h_idx + 1
        if body_idx >= len(doc.paragraphs):
            print(f"FAIL: Chapter {chapter_num} — no paragraph after Heading 1 at index {h_idx}")
            continue

        para = doc.paragraphs[body_idx]

        # Component: Check drop cap on this paragraph
        try:
            pPr = para._element.find(qn('w:pPr'))
            framePr = None
            if pPr is not None:
                framePr = pPr.find(qn('w:framePr'))

            if framePr is None:
                print(f"FAIL: Chapter {chapter_num} (para {body_idx}) — no framePr element (no drop cap)")
                continue

            drop_cap_val = framePr.get(qn('w:dropCap'))
            lines_val = framePr.get(qn('w:lines'))

            # Check dropCap attribute is 'drop' (or 'margin' which is also a valid drop cap type)
            if drop_cap_val not in ('drop', 'margin'):
                print(f"FAIL: Chapter {chapter_num} (para {body_idx}) — dropCap='{drop_cap_val}', expected 'drop'")
                continue

            # Check lines = '3'
            if lines_val == '3':
                # Full pass for this chapter
                print(f"PASS: Chapter {chapter_num} (para {body_idx}) — dropCap='{drop_cap_val}', lines='{lines_val}' ({points_per_chapter:.2f} pts)")
                total_score += points_per_chapter
            elif drop_cap_val in ('drop', 'margin'):
                # Partial: drop cap exists but wrong line count — give half credit
                print(f"PARTIAL: Chapter {chapter_num} (para {body_idx}) — lines='{lines_val}', expected '3' ({points_per_chapter * 0.5:.2f} pts)")
                total_score += points_per_chapter * 0.5
            else:
                print(f"FAIL: Chapter {chapter_num} (para {body_idx}) — dropCap='{drop_cap_val}', lines='{lines_val}'")

        except Exception as e:
            print(f"ERROR: Chapter {chapter_num} — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
