"""
Initial Setup: Enable drop caps for first paragraph of each chapter
Task ID: writer_fs_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_016'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Configure Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.size = Pt(18)

    # --- Chapter 1 ---
    doc.add_heading('Chapter 1: The Awakening', level=1)
    p = doc.add_paragraph(
        'Marina stood at the edge of the cliff, watching the sun dip below the horizon. '
        'The sea breeze carried the faint scent of salt and wildflowers, mingling with the '
        'distant cry of gulls wheeling overhead. She had returned to Thornfield after fifteen '
        'years, and everything felt simultaneously foreign and achingly familiar. The lighthouse '
        'still blinked its steady rhythm against the darkening sky, a beacon she had carried in '
        'her memory through all those years in the city.'
    )
    doc.add_paragraph(
        'The cottage behind her creaked in the wind, its weathered shutters tapping an irregular '
        'beat against the stone walls. Her grandmother had left it to her in the will, along with '
        'a leather-bound journal filled with cryptic entries and pressed flowers from gardens that '
        'no longer existed. Marina pulled her coat tighter and turned toward the narrow path '
        'leading down to the village.'
    )
    doc.add_paragraph(
        'By the time she reached the cobblestone main street, the lamplighter had already made '
        'his rounds. Warm golden light spilled from the windows of the Anchor and Crown, where '
        'laughter and the clink of glasses promised company she was not yet sure she wanted.'
    )

    # --- Chapter 2 ---
    doc.add_heading('Chapter 2: Echoes of the Past', level=1)
    p = doc.add_paragraph(
        'Beneath the floorboards of the cottage, Marina discovered a tin box wrapped in oilcloth. '
        'Inside lay a collection of letters dated from 1962 to 1974, written in a hand she did '
        'not recognise. The ink had faded to a pale sepia, but the words were still legible, '
        'filled with longing and references to a place called Briar Hollow that appeared on no '
        'modern map. She spread the letters across the kitchen table, arranging them by date, '
        'and began to read.'
    )
    doc.add_paragraph(
        'The first letter spoke of a summer storm that had uprooted the old oak beside the '
        'churchyard. The writer described sheltering in the bell tower, listening to the rain '
        'hammer against the slate roof while the bells swayed but never quite rang. There was '
        'a tenderness in the prose that suggested the recipient was deeply loved, though no '
        'name was ever used, only the salutation "My Dearest."'
    )
    doc.add_paragraph(
        'Marina set the letter down and stared out the kitchen window. The oak still stood by '
        'the churchyard, gnarled and enormous. Either the storm had never happened, or the tree '
        'had grown back with a persistence that bordered on the supernatural.'
    )

    # --- Chapter 3 ---
    doc.add_heading('Chapter 3: The Stranger at the Gate', level=1)
    p = doc.add_paragraph(
        'On the third morning, a man appeared at the garden gate. He was tall, with silver-grey '
        'hair cropped close to his skull and eyes the colour of winter slate. He carried a '
        'battered leather satchel over one shoulder and stood perfectly still, as though waiting '
        'for permission to exist on her threshold. Marina watched him through the kitchen curtain '
        'for a full minute before opening the door.'
    )
    doc.add_paragraph(
        '"You must be Eleanor\'s granddaughter," he said, his voice low and unhurried. "My name '
        'is Thomas Raith. I was a friend of your grandmother, many years ago. She asked me to '
        'give you something when you finally came back." He reached into the satchel and produced '
        'a small brass key, tarnished green with age, attached to a leather cord.'
    )
    doc.add_paragraph(
        'Marina took the key and turned it over in her palm. It was warm, as though it had been '
        'held for a long time. "What does it open?" she asked. Thomas Raith smiled, a slow, '
        'careful expression that did not quite reach his eyes. "That," he said, "is something '
        'you will need to discover for yourself."'
    )

    # --- Chapter 4 ---
    doc.add_heading('Chapter 4: The Hidden Room', level=1)
    p = doc.add_paragraph(
        'Following the instructions in her grandmother\'s journal, Marina descended the narrow '
        'staircase behind the pantry wall. The brass key fit perfectly into a lock concealed '
        'beneath a loose stone at the bottom step. Beyond the door lay a room she had never '
        'known existed: low-ceilinged, lined with bookshelves, and illuminated by a single '
        'porthole window that looked out onto the sea. Dust motes floated in the slanting light '
        'like tiny golden worlds suspended in amber.'
    )
    doc.add_paragraph(
        'The shelves held not only books but jars of preserved specimens, rolled nautical charts, '
        'and a collection of brass instruments whose purposes Marina could only guess at. In the '
        'centre of the room stood a writing desk, its surface covered in a fine layer of dust '
        'except for a single rectangle where something had recently been removed.'
    )
    doc.add_paragraph(
        'She ran her fingers along the spines of the books. Many were in languages she could not '
        'read. Others bore titles that seemed to shift when she looked at them from the corner '
        'of her eye. One slim volume, bound in deep blue leather, fell open in her hands to a '
        'page marked with a dried sprig of rosemary.'
    )

    # --- Chapter 5 ---
    doc.add_heading('Chapter 5: The Tide Turns', level=1)
    p = doc.add_paragraph(
        'That evening, the sea withdrew further than anyone in the village could remember. The '
        'harbour floor lay exposed, a landscape of glistening rock and stranded kelp, scattered '
        'with objects that had no business being there: a child\'s bicycle, a stone anchor carved '
        'with symbols, a glass bottle containing what appeared to be a miniature ship with sails '
        'of real silk. Marina walked among the debris with Thomas Raith at her side, both of them '
        'silent, both aware that something fundamental had shifted.'
    )
    doc.add_paragraph(
        '"It happens every century or so," Thomas said at last, bending to examine the carved '
        'anchor. "The sea gives back what it has taken. Your grandmother called it the Reckoning. '
        'She believed that Thornfield sits on a threshold, a place where the boundary between '
        'what is remembered and what is forgotten grows thin." He straightened and looked at her '
        'with those winter-slate eyes. "She also believed you would be the one to close it."'
    )
    doc.add_paragraph(
        'Marina stared at the retreating waterline, where phosphorescent light flickered and '
        'danced like foxfire. The brass key hung heavy around her neck. For the first time since '
        'arriving in Thornfield, she felt not the weight of inheritance but the pull of purpose, '
        'as though every road she had ever walked had been leading her here, to this beach, to '
        'this moment, to the edge of something vast and unknown.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
