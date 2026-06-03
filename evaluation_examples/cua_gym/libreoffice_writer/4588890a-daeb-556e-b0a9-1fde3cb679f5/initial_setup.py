"""
Initial Setup: Create a Writer document with meeting notes in plain text paragraphs.
Task ID: writer_lec_022
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_022'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Document title
    heading = doc.add_heading("Q1 2025 Product Strategy Meeting", level=1)

    # Meeting metadata
    meta = doc.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run("March 18, 2025")
    meta2 = doc.add_paragraph()
    meta2.add_run("Location: ").bold = True
    meta2.add_run("Conference Room B, 3rd Floor")
    meta3 = doc.add_paragraph()
    meta3.add_run("Attendees: ").bold = True
    meta3.add_run("Sarah Chen (VP Product), Marcus Johnson (Engineering Lead), "
                   "Priya Patel (UX Director), David Kim (QA Manager), "
                   "Elena Rodriguez (Marketing Lead)")

    doc.add_paragraph("")  # spacer

    # Section: Opening Remarks
    doc.add_heading("Opening Remarks", level=2)
    doc.add_paragraph(
        "Sarah Chen opened the meeting at 10:00 AM and welcomed the team. "
        "She emphasized the importance of aligning on Q1 priorities before the "
        "board review scheduled for April 3rd."
    )

    # Section: Product Roadmap Updates - plain text paragraphs (these will get the list style in golden)
    doc.add_heading("Key Discussion Points", level=2)

    meeting_notes = [
        "The mobile app redesign is on track for beta release by April 15th. "
        "Priya shared updated wireframes that received positive feedback from "
        "the focus group sessions conducted last week.",

        "Marcus reported that the backend migration to microservices architecture "
        "is 65% complete. The authentication service and user profile service have "
        "been fully migrated. Payment processing migration is scheduled for next sprint.",

        "David raised concerns about regression testing coverage for the new API "
        "endpoints. The team agreed to allocate two additional QA engineers from "
        "the platform team to assist with test automation.",

        "Elena presented the Q1 marketing campaign results showing a 23% increase "
        "in user acquisition compared to Q4 2024. The content marketing initiative "
        "drove 40% of new signups.",

        "The team discussed the customer feedback report highlighting requests for "
        "dark mode support, improved search functionality, and better notification "
        "management. These items will be prioritized in the Q2 roadmap.",

        "Sarah proposed establishing a weekly cross-functional sync meeting to "
        "improve coordination between engineering, design, and marketing teams "
        "during the critical pre-launch period.",

        "Budget allocation for Q2 was reviewed. Engineering will receive an "
        "additional $150,000 for cloud infrastructure scaling, and marketing "
        "will get $75,000 for the product launch campaign.",
    ]

    for note in meeting_notes:
        doc.add_paragraph(note)

    # Section: Action Items
    doc.add_heading("Action Items", level=2)
    doc.add_paragraph(
        "Priya to finalize the mobile app design specifications by March 25th "
        "and share with the engineering team for implementation review."
    )
    doc.add_paragraph(
        "Marcus to prepare a detailed migration timeline for the remaining "
        "microservices and present it at the next engineering standup."
    )
    doc.add_paragraph(
        "David to draft a comprehensive test plan for the new API endpoints "
        "and coordinate with the platform QA team by end of week."
    )
    doc.add_paragraph(
        "Elena to prepare the Q2 marketing strategy proposal incorporating "
        "the product launch timeline and submit for review by March 28th."
    )

    # Section: Next Meeting
    doc.add_heading("Next Meeting", level=2)
    doc.add_paragraph(
        "The next product strategy meeting is scheduled for April 1, 2025, "
        "at 10:00 AM in Conference Room B. Sarah requested that all team leads "
        "prepare department status updates in advance."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
