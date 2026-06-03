"""
Initial Setup: Thesis document with uniform Default Page Style (2.54cm margins)
Task ID: wrpara_040
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_040'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up the single section with Default Page Style margins (2.54cm = 1 inch)
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)  # A4
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ============================================================
    # PAGE 1: Title Page
    # ============================================================
    # Add some vertical spacing at the top
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)

    # University name
    uni = doc.add_paragraph()
    uni.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    uni.paragraph_format.space_after = Pt(24)
    run = uni.add_run("WESTFIELD UNIVERSITY")
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run.bold = True

    # Department
    dept = doc.add_paragraph()
    dept.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept.paragraph_format.space_after = Pt(48)
    run = dept.add_run("Department of Computer Science")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # Thesis title
    title = doc.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(36)
    run = title.add_run("Adaptive Neural Architectures for Real-Time\nSpeech Recognition in Noisy Environments")
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    run.bold = True

    # Author info
    for text in [
        "A thesis submitted in partial fulfillment",
        "of the requirements for the degree of",
        "Master of Science in Computer Science",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    # Spacer
    doc.add_paragraph()

    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_after = Pt(6)
    run = author.add_run("by")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    name = doc.add_paragraph()
    name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name.paragraph_format.space_after = Pt(24)
    run = name.add_run("Elena Vasquez Rodriguez")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.bold = True

    # Supervisor
    sup = doc.add_paragraph()
    sup.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sup.paragraph_format.space_after = Pt(6)
    run = sup.add_run("Supervisor: Prof. James Whitaker")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Date
    date_p = doc.add_paragraph()
    date_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_p.add_run("March 2026")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # ============================================================
    # PAGE 2: Abstract (use a page break to get to page 2)
    # ============================================================
    doc.add_page_break()

    abs_title = doc.add_paragraph()
    abs_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    abs_title.paragraph_format.space_after = Pt(18)
    run = abs_title.add_run("Abstract")
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run.bold = True

    abstract_text = (
        "This thesis presents a novel framework for adaptive neural architectures "
        "designed specifically for real-time speech recognition in acoustically "
        "challenging environments. Traditional speech recognition systems suffer "
        "significant performance degradation when exposed to non-stationary noise "
        "sources such as urban traffic, industrial machinery, and overlapping speech. "
        "Our proposed approach leverages dynamically reconfigurable neural network "
        "topologies that adjust their computational depth and feature extraction "
        "pathways based on the estimated signal-to-noise ratio of the input audio stream."
    )
    abs_p = doc.add_paragraph()
    abs_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abs_p.paragraph_format.line_spacing = 1.5
    run = abs_p.add_run(abstract_text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    abstract_text2 = (
        "Experimental results conducted across four benchmark datasets—LibriSpeech, "
        "CHiME-4, AURORA-4, and our custom WestfieldNoise corpus—demonstrate that "
        "the adaptive architecture achieves a word error rate (WER) improvement of "
        "12.3% to 18.7% over fixed-topology baselines in high-noise conditions "
        "(SNR below 5 dB), while maintaining comparable performance in clean audio "
        "scenarios. The system operates within a 150ms latency budget, making it "
        "suitable for deployment in real-time applications including voice assistants, "
        "hearing aids, and automated transcription services."
    )
    abs_p2 = doc.add_paragraph()
    abs_p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abs_p2.paragraph_format.line_spacing = 1.5
    run = abs_p2.add_run(abstract_text2)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_before = Pt(12)
    run = keywords.add_run("Keywords: ")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.bold = True
    run2 = keywords.add_run(
        "speech recognition, adaptive neural networks, noise robustness, "
        "real-time processing, dynamic architecture"
    )
    run2.font.size = Pt(12)
    run2.font.name = "Times New Roman"

    # ============================================================
    # PAGE 3+: Thesis Body (page break to start body)
    # ============================================================
    doc.add_page_break()

    # Chapter 1: Introduction
    ch1_title = doc.add_paragraph()
    ch1_title.paragraph_format.space_after = Pt(12)
    run = ch1_title.add_run("Chapter 1: Introduction")
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run.bold = True

    sec1_1 = doc.add_paragraph()
    sec1_1.paragraph_format.space_after = Pt(6)
    run = sec1_1.add_run("1.1  Background and Motivation")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.bold = True

    intro_text = (
        "The field of automatic speech recognition (ASR) has witnessed remarkable "
        "progress over the past decade, driven primarily by advances in deep learning "
        "architectures and the availability of large-scale training corpora. Modern "
        "ASR systems based on transformer models and end-to-end architectures routinely "
        "achieve word error rates below 5% on clean speech benchmarks such as "
        "LibriSpeech test-clean. However, these impressive results often fail to "
        "translate to real-world deployment scenarios where acoustic conditions are "
        "far from ideal."
    )
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(intro_text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    intro_text2 = (
        "Urban environments present particularly demanding acoustic conditions for "
        "speech recognition systems. A typical urban soundscape may include traffic "
        "noise with spectral energy concentrated below 1 kHz, construction machinery "
        "producing impulsive and broadband noise, competing speakers in crowded spaces, "
        "and reverberant conditions in enclosed public areas. The signal-to-noise ratio "
        "in such environments frequently drops below 0 dB, a regime where conventional "
        "ASR systems experience dramatic performance degradation."
    )
    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.first_line_indent = Cm(1.27)
    run = p2.add_run(intro_text2)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    sec1_2 = doc.add_paragraph()
    sec1_2.paragraph_format.space_before = Pt(12)
    sec1_2.paragraph_format.space_after = Pt(6)
    run = sec1_2.add_run("1.2  Research Objectives")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.bold = True

    objectives_text = (
        "This research addresses the fundamental limitation of fixed-topology neural "
        "networks in handling diverse acoustic conditions. We propose an adaptive "
        "architecture that dynamically adjusts its computational graph based on "
        "real-time noise estimation, allocating additional processing capacity to "
        "challenging audio segments while maintaining efficiency for cleaner signals. "
        "The primary objectives of this thesis are threefold: (1) to develop a "
        "differentiable architecture search mechanism suitable for streaming audio, "
        "(2) to design noise-aware gating modules that route audio features through "
        "condition-appropriate processing pathways, and (3) to validate the proposed "
        "system against established benchmarks under a strict real-time latency "
        "constraint of 150 milliseconds."
    )
    p3 = doc.add_paragraph()
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p3.paragraph_format.line_spacing = 1.5
    p3.paragraph_format.first_line_indent = Cm(1.27)
    run = p3.add_run(objectives_text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    sec1_3 = doc.add_paragraph()
    sec1_3.paragraph_format.space_before = Pt(12)
    sec1_3.paragraph_format.space_after = Pt(6)
    run = sec1_3.add_run("1.3  Thesis Organization")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.bold = True

    org_text = (
        "The remainder of this thesis is organized as follows. Chapter 2 provides a "
        "comprehensive review of related work spanning speech recognition, noise-robust "
        "feature extraction, and neural architecture search. Chapter 3 describes the "
        "proposed adaptive neural architecture in detail, including the noise estimation "
        "module, dynamic routing mechanism, and training procedure. Chapter 4 presents "
        "the experimental setup, including datasets, evaluation metrics, and baseline "
        "comparisons. Chapter 5 reports and analyzes the experimental results. Finally, "
        "Chapter 6 concludes the thesis and outlines directions for future work."
    )
    p4 = doc.add_paragraph()
    p4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p4.paragraph_format.line_spacing = 1.5
    p4.paragraph_format.first_line_indent = Cm(1.27)
    run = p4.add_run(org_text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Save the document - single section, uniform 2.54cm margins
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
