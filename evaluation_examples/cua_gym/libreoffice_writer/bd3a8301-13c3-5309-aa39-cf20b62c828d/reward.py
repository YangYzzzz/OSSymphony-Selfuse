"""
Reward Script: Roman numeral page numbering for front matter, Arabic for body
Task ID: writer_biz_063
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Document has multiple sections (section break exists)
  Component 2 (0.30): Front matter section uses lowerRoman page number format
  Component 3 (0.25): Body section uses decimal format restarting at 1
  Component 4 (0.25): Footer field codes reflect the numbering formats
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_063'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the document has Roman numeral page numbering for front matter
    and Arabic numerals starting at 1 for the body section.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    sections = list(doc.sections)
    num_sections = len(sections)

    # Component 1: Document has multiple sections (0.20 points)
    # Initial doc has 1 section; golden has 2+ (section break inserted)
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 -- Document has {num_sections} sections (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 -- Document has only {num_sections} section, expected >= 2")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Front matter section uses lowerRoman page number format (0.30 points)
    # Check section 0 (front matter) for pgNumType fmt=lowerRoman
    try:
        if num_sections >= 2:
            sec0 = sections[0]
            sectPr0 = sec0._sectPr
            pgNumType0 = sectPr0.find(qn('w:pgNumType'))
            if pgNumType0 is not None:
                fmt0 = pgNumType0.get(qn('w:fmt'))
                if fmt0 == 'lowerRoman':
                    print(f"PASS: Component 2 -- Front matter section has lowerRoman format (0.30 pts)")
                    total_score += 0.30
                else:
                    print(f"FAIL: Component 2 -- Front matter section pgNumType fmt='{fmt0}', expected 'lowerRoman'")
            else:
                print(f"FAIL: Component 2 -- Front matter section has no pgNumType element")
        else:
            print(f"FAIL: Component 2 -- Cannot check front matter format, only 1 section exists")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Body section uses decimal format restarting at page 1 (0.25 points)
    # Check section 1 (body) for pgNumType fmt=decimal and start=1
    try:
        if num_sections >= 2:
            sec1 = sections[1]
            sectPr1 = sec1._sectPr
            pgNumType1 = sectPr1.find(qn('w:pgNumType'))
            if pgNumType1 is not None:
                fmt1 = pgNumType1.get(qn('w:fmt'))
                start1 = pgNumType1.get(qn('w:start'))
                # fmt can be 'decimal' or None (default is decimal)
                fmt_ok = fmt1 in ('decimal', None)
                start_ok = start1 == '1'
                if fmt_ok and start_ok:
                    print(f"PASS: Component 3 -- Body section has decimal format starting at 1 (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 3 -- Body section pgNumType fmt='{fmt1}' start='{start1}', expected decimal/start=1")
            else:
                print(f"FAIL: Component 3 -- Body section has no pgNumType element (no restart)")
        else:
            print(f"FAIL: Component 3 -- Cannot check body format, only 1 section exists")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Footer field codes reflect numbering formats (0.25 points)
    # Front matter footer should contain roman formatting; body footer should have PAGE field
    # This verifies the footers are properly set up (not linked to previous, correct instrText)
    try:
        if num_sections >= 2:
            sub_score = 0.0

            # Check section 0 footer: should have PAGE field with roman switch or be unlinked
            sec0_footer = sections[0].footer
            sec0_footer_linked = sections[0].footer.is_linked_to_previous

            # Extract instrText from section 0 footer
            sec0_instr = []
            for fp in sec0_footer.paragraphs:
                for run in fp.runs:
                    for it in run.element.findall(qn('w:instrText')):
                        sec0_instr.append(it.text)

            # Check section 0 has a PAGE field (with or without roman switch)
            sec0_has_page = any('PAGE' in t for t in sec0_instr)
            sec0_has_roman = any('roman' in t.lower() for t in sec0_instr)

            if sec0_has_page and sec0_has_roman:
                print(f"PASS: Component 4a -- Front matter footer has PAGE with roman switch (0.125 pts)")
                sub_score += 0.125
            elif sec0_has_page:
                # Has PAGE but no roman switch -- partial (the pgNumType handles formatting)
                print(f"PARTIAL: Component 4a -- Front matter footer has PAGE field but no roman switch (0.0625 pts)")
                sub_score += 0.0625
            else:
                print(f"FAIL: Component 4a -- Front matter footer missing PAGE field")

            # Check section 1 footer: should not be linked to section 0 and should have PAGE field
            sec1_footer = sections[1].footer
            sec1_footer_linked = sections[1].footer.is_linked_to_previous

            sec1_instr = []
            for fp in sec1_footer.paragraphs:
                for run in fp.runs:
                    for it in run.element.findall(qn('w:instrText')):
                        sec1_instr.append(it.text)

            sec1_has_page = any('PAGE' in t for t in sec1_instr)

            if sec1_has_page and not sec1_footer_linked:
                print(f"PASS: Component 4b -- Body footer has PAGE field and is unlinked (0.125 pts)")
                sub_score += 0.125
            elif sec1_has_page:
                print(f"PARTIAL: Component 4b -- Body footer has PAGE field but is linked to previous (0.0625 pts)")
                sub_score += 0.0625
            else:
                print(f"FAIL: Component 4b -- Body footer missing PAGE field or linked to previous")

            total_score += sub_score
            print(f"  Component 4 total: {sub_score} pts")
        else:
            print(f"FAIL: Component 4 -- Cannot check footers, only 1 section exists")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(round(total_score, 4), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
