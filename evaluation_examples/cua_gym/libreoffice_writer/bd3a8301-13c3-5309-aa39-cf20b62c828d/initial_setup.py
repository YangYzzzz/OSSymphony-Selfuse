"""
Initial Setup: Proposal document with front matter and body content, all using Arabic page numbering.
Task ID: writer_biz_063
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_063'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph."""
    run1 = paragraph.add_run()
    fldChar1 = run1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run1._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = run2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._element.append(fldChar2)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # -- Footer with page number (Arabic, default) --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp)

    # ===== FRONT MATTER =====

    # -- Page 1: Table of Contents --
    heading = doc.add_heading('Table of Contents', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    toc_entries = [
        ('Executive Summary', '2'),
        ('1. Introduction', '4'),
        ('2. Market Analysis', '5'),
        ('3. Proposed Solution', '7'),
        ('4. Implementation Timeline', '9'),
        ('5. Budget and Resources', '10'),
        ('6. Risk Assessment', '11'),
        ('7. Conclusion', '12'),
    ]

    for entry, page in toc_entries:
        para = doc.add_paragraph()
        run = para.add_run(f'{entry}')
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        tab_run = para.add_run(f'\t{page}')
        tab_run.font.size = Pt(12)
        tab_run.font.name = 'Calibri'
        para.paragraph_format.space_after = Pt(4)

    # Page break to Executive Summary
    doc.add_page_break()

    # -- Pages 2-3: Executive Summary --
    heading = doc.add_heading('Executive Summary', level=1)

    para = doc.add_paragraph()
    run = para.add_run(
        'Meridian Technologies is pleased to present this comprehensive proposal for the '
        'Digital Transformation Initiative at Westfield Manufacturing Group. This document '
        'outlines our strategic approach to modernizing your core business operations through '
        'integrated technology solutions that will drive efficiency, reduce costs, and position '
        'your organization for sustained growth in an increasingly competitive landscape.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'Our analysis of Westfield\'s current operational infrastructure reveals significant '
        'opportunities for improvement across supply chain management, customer relationship '
        'management, and internal collaboration platforms. The proposed solution leverages '
        'cloud-native architectures and artificial intelligence to automate manual processes '
        'that currently consume an estimated 2,400 person-hours per quarter.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'Key deliverables include: (1) an enterprise resource planning system integrating '
        'procurement, inventory, and fulfillment; (2) a customer data platform unifying '
        'touchpoints across retail, wholesale, and e-commerce channels; (3) a real-time '
        'analytics dashboard providing actionable insights to executive leadership; and '
        '(4) comprehensive training programs ensuring seamless adoption across all departments.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'The total investment for this three-phase initiative is $4.2 million over 18 months, '
        'with projected annual savings of $1.8 million beginning in Year 2. Our conservative '
        'ROI model indicates break-even within 28 months and a five-year net present value '
        'of $3.7 million at a discount rate of 8%. These projections are based on benchmark '
        'data from comparable implementations at Fortune 500 manufacturers.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'Meridian Technologies brings over 15 years of experience delivering enterprise-scale '
        'digital transformations. Our team of 340 certified professionals has successfully '
        'completed 127 major deployments across the manufacturing sector, maintaining a 94% '
        'on-time delivery record and a client satisfaction rating of 4.8 out of 5.0. We are '
        'confident that this partnership will deliver transformative results for Westfield '
        'Manufacturing Group.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    # Page break to body
    doc.add_page_break()

    # ===== BODY CONTENT =====

    # -- Chapter 1: Introduction --
    heading = doc.add_heading('1. Introduction', level=1)

    para = doc.add_paragraph()
    run = para.add_run(
        'The manufacturing industry is undergoing a fundamental shift driven by Industry 4.0 '
        'technologies, changing consumer expectations, and global supply chain disruptions. '
        'Organizations that fail to adapt risk losing market share to more agile competitors '
        'who have embraced digital-first operational models.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'Westfield Manufacturing Group, founded in 1987 and headquartered in Portland, Oregon, '
        'has grown to become a leading provider of precision-engineered components for the '
        'automotive and aerospace sectors. With annual revenues of $890 million and operations '
        'spanning 12 facilities across North America, Westfield has built a reputation for '
        'quality and reliability that is the envy of the industry.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'However, recent assessments have identified critical gaps in Westfield\'s technology '
        'infrastructure that threaten to undermine its competitive position. Legacy systems, '
        'some dating back to 2005, are creating data silos, increasing maintenance costs, and '
        'limiting the organization\'s ability to respond quickly to market changes. This proposal '
        'addresses these challenges with a phased, pragmatic approach to digital transformation.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    # Page break
    doc.add_page_break()

    # -- Chapter 2: Market Analysis --
    heading = doc.add_heading('2. Market Analysis', level=1)

    para = doc.add_paragraph()
    run = para.add_run(
        'The global precision manufacturing market is projected to reach $412 billion by 2028, '
        'growing at a compound annual growth rate of 6.3%. Key growth drivers include the '
        'expansion of electric vehicle production, increased defense spending, and the reshoring '
        'of critical manufacturing capabilities to North America.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'Competitive benchmarking reveals that 78% of top-quartile manufacturers have already '
        'implemented or are actively deploying ERP modernization programs. Among Westfield\'s '
        'direct competitors, Apex Precision Corp completed a $6.1 million digital overhaul in '
        '2024, resulting in a 22% reduction in order-to-delivery cycle times. Similarly, '
        'NovaTech Industries reported a 31% improvement in inventory accuracy after migrating '
        'to a cloud-based supply chain platform.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run(
        'Customer expectations are evolving rapidly. A recent survey of 500 procurement managers '
        'in the automotive sector found that 89% now require real-time order tracking, 72% '
        'expect automated quality certification documentation, and 64% prefer suppliers who '
        'offer integrated digital collaboration portals. Westfield\'s current systems support '
        'none of these capabilities at scale.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    # Page break
    doc.add_page_break()

    # -- Chapter 3: Proposed Solution --
    heading = doc.add_heading('3. Proposed Solution', level=1)

    para = doc.add_paragraph()
    run = para.add_run(
        'Our proposed solution is built on three interconnected pillars: operational integration, '
        'data intelligence, and workforce enablement. Each pillar addresses specific pain points '
        'identified during our discovery phase while contributing to a cohesive digital ecosystem '
        'that will serve Westfield for the next decade and beyond.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    sub = doc.add_heading('3.1 Operational Integration', level=2)

    para = doc.add_paragraph()
    run = para.add_run(
        'The centerpiece of this pillar is a unified ERP platform based on SAP S/4HANA Cloud, '
        'customized for Westfield\'s multi-facility operations. This platform will replace the '
        'current patchwork of Oracle E-Business Suite (procurement), JD Edwards (inventory), '
        'and custom Access databases (production scheduling) with a single source of truth.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    sub = doc.add_heading('3.2 Data Intelligence', level=2)

    para = doc.add_paragraph()
    run = para.add_run(
        'A Microsoft Azure-based data lake will aggregate operational data from all facilities, '
        'customer interactions, and external market feeds. Machine learning models will provide '
        'predictive maintenance alerts, demand forecasting, and automated quality anomaly '
        'detection. Executive dashboards built on Power BI will deliver real-time KPIs to '
        'leadership at all levels of the organization.'
    )
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(8)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
