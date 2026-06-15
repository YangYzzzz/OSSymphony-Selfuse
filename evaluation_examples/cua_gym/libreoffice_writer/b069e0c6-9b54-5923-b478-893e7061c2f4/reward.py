"""
Reward Script: Environmental Impact Assessment Summary in LibreOffice Writer
Task ID: writer_wf_060
Domain: libreoffice_writer
Scoring:
  C1: Title text correct (0.10)
  C2: 7 Heading 1 sections with correct names (0.25)
  C3: TOC presence (0.10)
  C4: Impact Assessment table — 7 rows x 3 cols (0.20)
  C5: 5 numbered mitigation measures (0.10)
  C6: Monitoring Program table — 5 rows x 4 cols (0.15)
  C7: Double spacing on all paragraphs (0.10)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_060'

def persist_app_state(domain: str):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather basic info
    paragraphs = doc.paragraphs
    tables = doc.tables

    # Component 1: Title text correct (0.10 points)
    try:
        title_paras = [p for p in paragraphs if p.style and p.style.name == 'Title']
        if title_paras and 'environmental impact assessment' in title_paras[0].text.lower() and 'solar farm' in title_paras[0].text.lower():
            print(f"PASS: Component 1 — Title found: '{title_paras[0].text}' (0.10 pts)")
            total_score += 0.10
        else:
            found_text = title_paras[0].text if title_paras else 'NO TITLE PARA'
            print(f"FAIL: Component 1 — Expected EIA Solar Farm title, found: '{found_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 7 Heading 1 sections with correct names (0.25 points)
    try:
        expected_headings = [
            'project description',
            'baseline conditions',
            'impact assessment',
            'cumulative effects',
            'mitigation plan',
            'monitoring program',
            'conclusion'
        ]
        h1_paras = [p.text.strip().lower() for p in paragraphs if p.style and p.style.name == 'Heading 1']
        matched = 0
        for eh in expected_headings:
            if any(eh in h1.lower() for h1 in h1_paras):
                matched += 1
        if matched == 7:
            print(f"PASS: Component 2 — All 7 Heading 1 sections found (0.25 pts)")
            total_score += 0.25
        elif matched >= 5:
            partial = round(0.25 * (matched / 7), 2)
            print(f"PARTIAL: Component 2 — {matched}/7 headings found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {matched}/7 expected Heading 1 sections found. Found: {h1_paras}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: TOC presence (0.10 points)
    # Check for TOC text or field code in document
    try:
        has_toc = False
        # Check for TOC-related text in paragraphs
        for p in paragraphs:
            if 'table of contents' in p.text.lower():
                has_toc = True
                break
        # Also check XML for TOC field code
        if not has_toc:
            body_xml = doc.element.body.xml
            if 'TOC' in body_xml:
                has_toc = True
        if has_toc:
            print(f"PASS: Component 3 — Table of Contents found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No Table of Contents detected")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Impact Assessment table — header + 6 data rows, 3 columns (0.20 points)
    try:
        impact_table_found = False
        for t in tables:
            # Check if this is the impact assessment table by looking at headers
            if len(t.rows) >= 1 and len(t.columns) >= 3:
                header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
                if ('environmental factor' in header_cells[0] and
                    'impact level' in header_cells[1] and
                    'mitigation' in header_cells[2]):
                    # This is the impact assessment table
                    data_rows = len(t.rows) - 1  # minus header
                    col_count = len(t.columns)
                    if data_rows >= 6 and col_count >= 3:
                        print(f"PASS: Component 4 — Impact table: {data_rows} data rows x {col_count} cols (0.20 pts)")
                        total_score += 0.20
                        impact_table_found = True
                    elif data_rows >= 4:
                        partial = round(0.20 * (data_rows / 6), 2)
                        print(f"PARTIAL: Component 4 — Impact table has {data_rows}/6 data rows ({partial} pts)")
                        total_score += partial
                        impact_table_found = True
                    else:
                        print(f"FAIL: Component 4 — Impact table only has {data_rows} data rows (need >= 6)")
                        impact_table_found = True
                    break
        if not impact_table_found:
            print(f"FAIL: Component 4 — No impact assessment table found with correct headers")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: 5 numbered mitigation measures (0.10 points)
    try:
        # Find paragraphs under the "Mitigation Plan" heading that start with numbers
        in_mitigation = False
        numbered_measures = []
        for p in paragraphs:
            if p.style and p.style.name == 'Heading 1' and 'mitigation' in p.text.lower():
                in_mitigation = True
                continue
            if p.style and p.style.name == 'Heading 1' and in_mitigation:
                in_mitigation = False
                break
            if in_mitigation and p.text.strip():
                text = p.text.strip()
                # Check if starts with a number (1., 2., etc.) or is a list style
                if (text[0].isdigit() or
                    (p.style and 'list' in p.style.name.lower())):
                    numbered_measures.append(text[:50])

        if len(numbered_measures) >= 5:
            print(f"PASS: Component 5 — {len(numbered_measures)} mitigation measures found (0.10 pts)")
            total_score += 0.10
        elif len(numbered_measures) >= 3:
            partial = round(0.10 * (len(numbered_measures) / 5), 2)
            print(f"PARTIAL: Component 5 — {len(numbered_measures)}/5 measures ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {len(numbered_measures)} mitigation measures found (need 5)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Monitoring Program table — header + 4 data rows, 4 columns (0.15 points)
    try:
        monitoring_table_found = False
        for t in tables:
            if len(t.rows) >= 1 and len(t.columns) >= 4:
                header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
                if ('parameter' in header_cells[0] and
                    'frequency' in header_cells[1] and
                    'method' in header_cells[2] and
                    'responsible' in header_cells[3]):
                    data_rows = len(t.rows) - 1
                    col_count = len(t.columns)
                    if data_rows >= 4 and col_count >= 4:
                        print(f"PASS: Component 6 — Monitoring table: {data_rows} data rows x {col_count} cols (0.15 pts)")
                        total_score += 0.15
                        monitoring_table_found = True
                    elif data_rows >= 2:
                        partial = round(0.15 * (data_rows / 4), 2)
                        print(f"PARTIAL: Component 6 — Monitoring table has {data_rows}/4 data rows ({partial} pts)")
                        total_score += partial
                        monitoring_table_found = True
                    else:
                        print(f"FAIL: Component 6 — Monitoring table only has {data_rows} data rows (need >= 4)")
                        monitoring_table_found = True
                    break
        if not monitoring_table_found:
            print(f"FAIL: Component 6 — No monitoring program table found with correct headers")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Double spacing throughout (0.10 points)
    try:
        # Check that all non-empty paragraphs have line_spacing == 2.0
        total_checked = 0
        double_spaced = 0
        for p in paragraphs:
            if p.text.strip():
                total_checked += 1
                ls = p.paragraph_format.line_spacing
                if ls is not None and abs(float(ls) - 2.0) < 0.01:
                    double_spaced += 1

        if total_checked > 0:
            ratio = double_spaced / total_checked
            if ratio >= 0.9:
                print(f"PASS: Component 7 — {double_spaced}/{total_checked} paragraphs double-spaced (0.10 pts)")
                total_score += 0.10
            elif ratio >= 0.5:
                partial = round(0.10 * ratio, 2)
                print(f"PARTIAL: Component 7 — {double_spaced}/{total_checked} double-spaced ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 7 — Only {double_spaced}/{total_checked} paragraphs are double-spaced")
        else:
            print(f"FAIL: Component 7 — No non-empty paragraphs found")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
