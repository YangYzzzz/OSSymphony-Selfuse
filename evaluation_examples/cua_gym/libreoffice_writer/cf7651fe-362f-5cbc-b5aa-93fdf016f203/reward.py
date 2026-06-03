"""
Reward Script: Grant Application Summary in LibreOffice Writer
Task ID: writer_wf_029
Domain: libreoffice_writer
Scoring:
  C1 (0.10) - Title paragraph with correct text
  C2 (0.15) - Six Heading 1 sections
  C3 (0.15) - Four numbered list objectives
  C4 (0.15) - Budget table structure (header + 5 categories, 4 columns)
  C5 (0.15) - Budget total sums to $150,000
  C6 (0.15) - Timeline table (4 quarter rows, 2 columns)
  C7 (0.15) - Left margin set to 3cm for binding
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_029'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice edits via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.shared import Cm, Emu
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Title paragraph with correct text (0.10 points)
    try:
        title_found = False
        for para in doc.paragraphs:
            if para.style.name == 'Title' and 'grant application' in para.text.lower() and 'community health initiative' in para.text.lower():
                title_found = True
                break
        if title_found:
            print(f"PASS: Component 1 — Title found: '{para.text}' (0.10 pts)")
            total_score += 0.10
        else:
            # Also accept heading level 0
            for para in doc.paragraphs:
                if para.style.name == 'Heading 0' and 'grant application' in para.text.lower():
                    title_found = True
                    break
            if title_found:
                print(f"PASS: Component 1 — Title (Heading 0) found (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 1 — Title 'Grant Application - Community Health Initiative' not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Six Heading 1 sections (0.15 points)
    # Expected: Project Summary, Objectives, Target Population, Methodology, Evaluation Plan, Budget, Timeline
    # That's actually 7 but the task says 6 sections; Budget and Timeline are also H1 = total at least 6
    try:
        heading1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
        heading1_texts = [p.text.strip().lower() for p in heading1_paras]
        expected_sections = ['project summary', 'objectives', 'target population',
                             'methodology', 'evaluation plan', 'budget', 'timeline']
        # Count how many expected sections are present
        found_sections = []
        for exp in expected_sections:
            for ht in heading1_texts:
                if exp in ht:
                    found_sections.append(exp)
                    break

        # Need at least 6 of the 7 section headings (task says "6 sections" but lists 7 including Budget and Timeline)
        if len(found_sections) >= 6:
            print(f"PASS: Component 2 — {len(found_sections)} Heading 1 sections found: {found_sections} (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — Expected >= 6 Heading 1 sections, found {len(found_sections)}: {found_sections}")
            print(f"       All Heading 1 texts: {heading1_texts}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Four numbered list objectives (0.15 points)
    try:
        numbered_items = [p for p in doc.paragraphs if p.style.name in ('List Number', 'List Number 1', 'List Number 2')]
        if len(numbered_items) >= 4:
            print(f"PASS: Component 3 — {len(numbered_items)} numbered list items found (0.15 pts)")
            total_score += 0.15
        else:
            # Fallback: check if there are paragraphs that start with "1." etc. near the Objectives heading
            fallback_count = 0
            in_objectives = False
            for p in doc.paragraphs:
                if p.style.name == 'Heading 1' and 'objectives' in p.text.lower():
                    in_objectives = True
                    continue
                if in_objectives and p.style.name == 'Heading 1':
                    break
                if in_objectives and re.match(r'^\d+[\.\)]\s', p.text.strip()):
                    fallback_count += 1
            if fallback_count >= 4:
                print(f"PASS: Component 3 — {fallback_count} numbered items found (fallback) (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 — Expected 4 numbered objectives, found {len(numbered_items)} List Number + {fallback_count} manual numbering")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Budget table structure (0.15 points)
    # Expected: header row (Category, Year 1, Year 2, Total) + 5 category rows, 4 columns
    try:
        budget_table = None
        for t in doc.tables:
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            if 'category' in header_cells and ('year 1' in header_cells or 'year1' in header_cells):
                budget_table = t
                break

        if budget_table is not None:
            num_cols = len(budget_table.columns)
            # Count data rows (excluding header and any total row)
            data_rows = 0
            for ri, row in enumerate(budget_table.rows):
                if ri == 0:
                    continue  # skip header
                first_cell = row.cells[0].text.strip().lower()
                if first_cell != 'total' and first_cell != '':
                    data_rows += 1

            if num_cols >= 4 and data_rows >= 5:
                print(f"PASS: Component 4 — Budget table has {num_cols} cols, {data_rows} data rows (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — Budget table: {num_cols} cols (need 4), {data_rows} data rows (need 5)")
        else:
            print(f"FAIL: Component 4 — Budget table with expected headers not found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Budget total sums to $150,000 (0.15 points)
    try:
        if budget_table is not None:
            # Find the total row or last column values
            total_found = False
            for row in budget_table.rows:
                cells = [c.text.strip() for c in row.cells]
                if cells[0].lower() == 'total':
                    # Check the last cell (Total column) for $150,000
                    total_val = cells[-1].replace(',', '').replace('$', '').replace(' ', '')
                    if '150000' in total_val:
                        total_found = True
                        print(f"PASS: Component 5 — Budget total is $150,000 ({cells[-1]}) (0.15 pts)")
                        total_score += 0.15
                    else:
                        print(f"FAIL: Component 5 — Budget total is '{cells[-1]}', expected $150,000")
                    break

            if not total_found:
                # Sum all "Total" column values for data rows
                total_col_idx = None
                header_cells = [c.text.strip().lower() for c in budget_table.rows[0].cells]
                for idx, h in enumerate(header_cells):
                    if h == 'total':
                        total_col_idx = idx
                        break

                if total_col_idx is not None:
                    running_sum = 0
                    for ri, row in enumerate(budget_table.rows):
                        if ri == 0:
                            continue
                        val = row.cells[total_col_idx].text.strip().replace(',', '').replace('$', '').replace(' ', '')
                        try:
                            running_sum += int(val)
                        except ValueError:
                            pass
                    if running_sum == 150000:
                        print(f"PASS: Component 5 — Budget category totals sum to $150,000 (0.15 pts)")
                        total_score += 0.15
                    else:
                        print(f"FAIL: Component 5 — Budget category totals sum to ${running_sum:,}, expected $150,000")
                else:
                    print(f"FAIL: Component 5 — Could not find 'Total' column in budget table")
        else:
            print(f"FAIL: Component 5 — No budget table found to check total")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Timeline table (4 quarter rows, 2 columns) (0.15 points)
    try:
        timeline_table = None
        for t in doc.tables:
            header_cells = [c.text.strip().lower() for c in t.rows[0].cells]
            if ('quarter' in header_cells or any('quarter' in h for h in header_cells)) and \
               ('activities' in header_cells or any('activit' in h for h in header_cells)):
                timeline_table = t
                break

        if timeline_table is not None:
            num_cols = len(timeline_table.columns)
            # Data rows (excluding header)
            data_rows = len(timeline_table.rows) - 1
            quarter_rows = 0
            for ri, row in enumerate(timeline_table.rows):
                if ri == 0:
                    continue
                first_cell = row.cells[0].text.strip().lower()
                if 'q' in first_cell or 'quarter' in first_cell:
                    quarter_rows += 1

            if num_cols >= 2 and quarter_rows >= 4:
                print(f"PASS: Component 6 — Timeline table has {num_cols} cols, {quarter_rows} quarter rows (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 6 — Timeline table: {num_cols} cols (need 2), {quarter_rows} quarter rows (need 4)")
        else:
            print(f"FAIL: Component 6 — Timeline table with Quarter/Activities headers not found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    # Component 7: Left margin set to 3cm for binding (0.15 points)
    try:
        section = doc.sections[0]
        left_margin_cm = section.left_margin / 360000.0  # EMU to cm
        # 3cm = 1080000 EMU. Allow tolerance of 0.15cm (~0.06 inches)
        if abs(left_margin_cm - 3.0) <= 0.15:
            print(f"PASS: Component 7 — Left margin is {left_margin_cm:.2f} cm (target 3.0 cm) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 7 — Left margin is {left_margin_cm:.2f} cm, expected ~3.0 cm")
    except Exception as e:
        print(f"ERROR: Component 7 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entrypoint
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
