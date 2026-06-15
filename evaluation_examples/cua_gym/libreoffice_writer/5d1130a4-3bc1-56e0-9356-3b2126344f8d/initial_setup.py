"""
Initial Setup: Create project_tracker.xlsx and project_report.docx for project status update task
Task ID: osworld_multi_apps_calc_to_writer_007
Domain: libreoffice_writer (multi-app: Writer open, Calc file in ~/Documents/)
"""

import os
import shlex
import subprocess
import time
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
DOCS_DIR = '/home/user/Documents'
TASK_ID = 'osworld_multi_apps_calc_to_writer_007'
TRACKER_PATH = f'{DOCS_DIR}/project_tracker.xlsx'
REPORT_PATH = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_tracker():
    """Create ~/Documents/project_tracker.xlsx with mixed-status project tasks."""
    os.makedirs(DOCS_DIR, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Project Tasks'

    # Header row styling
    header_font = Font(bold=True, size=12, color='FFFFFF')
    header_fill = PatternFill(start_color='FF2F5496', end_color='FF2F5496', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center')
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    headers = ['Task', 'Owner', 'Deadline', 'Status']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = border

    # Project task data — 5 Completed, 2 In Progress, 8 Not Started
    # Columns: Task, Owner, Deadline, Status
    tasks = [
        ('Requirements Gathering',      'Emily Torres',     '2025-01-15', 'Completed'),
        ('System Architecture Design',  'James Whitfield',  '2025-01-30', 'Completed'),
        ('Database Schema Draft',       'Priya Nair',       '2025-02-10', 'Completed'),
        ('API Endpoint Specification',  'Marco Delgado',    '2025-02-20', 'Completed'),
        ('Frontend Prototype',          'Sarah Chen',       '2025-03-05', 'Completed'),
        ('Backend Core Module',         "Liam O'Brien",     '2025-03-25', 'In Progress'),
        ('Unit Test Suite',             'Ayesha Rahman',    '2025-04-10', 'In Progress'),
        ('Integration Testing',         'Derek Nguyen',     '2025-04-25', 'Not Started'),
        ('Security Audit',              'Fatima Al-Hassan', '2025-05-05', 'Not Started'),
        ('Performance Benchmarking',    'Marcus Johnson',   '2025-05-15', 'Not Started'),
        ('User Acceptance Testing',     'Clara Petrov',     '2025-05-30', 'Not Started'),
        ('Deployment Pipeline Setup',   'Hiroshi Tanaka',   '2025-06-10', 'Not Started'),
        ('Documentation Review',        'Amara Osei',       '2025-06-20', 'Not Started'),
        ('Stakeholder Demo Preparation','Emily Torres',     '2025-06-28', 'Not Started'),
        ('Go-Live Release',             'James Whitfield',  '2025-07-01', 'Not Started'),
    ]

    for r, (task, owner, deadline, status) in enumerate(tasks, 2):
        for col, val in enumerate([task, owner, deadline, status], 1):
            cell = ws.cell(row=r, column=col, value=val)
            cell.border = border
        # Color-code status column
        status_cell = ws.cell(row=r, column=4)
        if status == 'Completed':
            status_cell.font = Font(color='FF006400')  # dark green
        elif status == 'In Progress':
            status_cell.font = Font(color='FF8B4513')  # amber/brown

    # Column widths
    ws.column_dimensions['A'].width = 32
    ws.column_dimensions['B'].width = 22
    ws.column_dimensions['C'].width = 14
    ws.column_dimensions['D'].width = 14
    ws.row_dimensions[1].height = 22

    wb.save(TRACKER_PATH)
    print(f'Tracker created: {TRACKER_PATH}')


def create_report():
    """Create project_report.docx with a Completed Milestones section (NO table yet)."""
    doc = Document()

    # Document title
    title = doc.add_heading('Project Status Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Meta line
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Q2 2025  |  Project Phoenix  |  Prepared by: PMO Team')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph('')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'Project Phoenix is progressing on schedule with five key milestones successfully '
        'delivered in Q1 2025. The core backend development is currently underway, with '
        'integration testing and deployment pipeline work scheduled for Q2. Stakeholder '
        'confidence remains high following the successful frontend prototype demonstration.'
    )
    doc.add_paragraph('')

    # Project Overview
    doc.add_heading('Project Overview', level=1)
    doc.add_paragraph(
        'Project Phoenix aims to deliver a next-generation customer relationship management '
        'platform with AI-assisted analytics, real-time reporting dashboards, and seamless '
        'third-party integrations. The project team consists of fifteen specialists across '
        'engineering, design, QA, and product management.'
    )
    for item in [
        'Project Start Date: November 15, 2024',
        'Planned Go-Live: July 1, 2025',
        'Total Budget: $1,450,000',
        'Sponsor: VP of Technology, Richard Harmon',
        'Project Manager: Emily Torres',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph('')

    # Current Status
    doc.add_heading('Current Status', level=1)
    doc.add_paragraph(
        'As of the reporting date, the project is 33% complete by milestone count. '
        'Five out of fifteen planned milestones have been delivered. Two milestones '
        'are currently in progress, with the remaining eight scheduled for subsequent phases.'
    )
    doc.add_paragraph('')

    # Completed Milestones — agent must add a table here from the spreadsheet
    doc.add_heading('Completed Milestones', level=1)
    doc.add_paragraph(
        'The following milestones have been completed and formally signed off by the '
        'project sponsor. All deliverables are archived in the project repository.'
    )
    # NOTE: No table here — the agent must extract Status=Completed rows from
    # ~/Documents/project_tracker.xlsx and insert a table in this section.
    doc.add_paragraph('')

    # Upcoming Milestones
    doc.add_heading('Upcoming Milestones', level=1)
    doc.add_paragraph(
        'The following key milestones are scheduled for completion in the coming weeks. '
        'The team is on track to meet all upcoming deadlines barring any unforeseen blockers.'
    )
    for item in [
        'Backend Core Module — due March 25, 2025 (In Progress)',
        'Unit Test Suite — due April 10, 2025 (In Progress)',
        'Integration Testing — due April 25, 2025',
        'Security Audit — due May 5, 2025',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph('')

    # Risks and Issues
    doc.add_heading('Risks and Issues', level=1)
    doc.add_paragraph(
        'The following risks have been identified and are being actively monitored.'
    )
    for risk in [
        'Third-party API dependency delays — Medium risk, mitigation plan in place',
        'Resource availability during UAT phase — Low risk, contingency resources identified',
        'Scope creep from stakeholder feedback — Medium risk, change control board active',
    ]:
        doc.add_paragraph(risk, style='List Bullet')

    doc.save(REPORT_PATH)
    print(f'Report created: {REPORT_PATH}')


def main():
    create_tracker()
    create_report()

    # GUI-ready startup: open the report in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{REPORT_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with project report (DISPLAY=:0)')


main()
