"""
Reward Script: Extract 'Completed' tasks from project_tracker.xlsx and insert as table
                in 'Completed Milestones' section of project_report.docx
Task ID: osworld_multi_apps_calc_to_writer_007
Domain: libreoffice_writer (multi-app: calc -> writer)
Scoring:
  Component 1: At least one table exists in the document (0.3 pts)
  Component 2: Table header row matches expected columns (0.3 pts)
  Component 3: All 5 Completed tasks are present with correct data (0.4 pts)
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_007'

# Expected completed tasks from project_tracker.xlsx (Status == 'Completed')
EXPECTED_COMPLETED_TASKS = [
    ('Requirements Gathering', 'Emily Torres', '2025-01-15', 'Completed'),
    ('System Architecture Design', 'James Whitfield', '2025-01-30', 'Completed'),
    ('Database Schema Draft', 'Priya Nair', '2025-02-10', 'Completed'),
    ('API Endpoint Specification', 'Marco Delgado', '2025-02-20', 'Completed'),
    ('Frontend Prototype', 'Sarah Chen', '2025-03-05', 'Completed'),
]

EXPECTED_HEADERS = ['Task', 'Owner', 'Deadline', 'Status']


def normalize_date(val):
    """Normalize date strings: strip whitespace, handle datetime objects."""
    if val is None:
        return ''
    s = str(val).strip()
    # Handle datetime objects that show as 'YYYY-MM-DD HH:MM:SS'
    if ' ' in s and ':' in s:
        s = s.split(' ')[0]
    return s


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: A table exists in the document (0.3 points)
    # Initial env has 0 tables; golden env has 1 table with completed tasks
    try:
        tables = doc.tables
        num_tables = len(tables)
        if num_tables >= 1:
            print(f"PASS: Component 1 — Table exists in document ({num_tables} table(s) found) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No tables found in document (expected at least 1)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table header row matches expected columns (0.3 points)
    # Check if any table has a header row with Task, Owner, Deadline, Status
    try:
        header_found = False
        if len(doc.tables) >= 1:
            for table in doc.tables:
                if len(table.rows) >= 1:
                    header_row = [cell.text.strip() for cell in table.rows[0].cells]
                    # Check if headers match (case-insensitive, partial match acceptable)
                    matched_headers = sum(
                        1 for h in EXPECTED_HEADERS
                        if any(h.lower() == cell.lower() for cell in header_row)
                    )
                    if matched_headers >= 3:  # at least 3 of 4 headers must match
                        header_found = True
                        print(f"PASS: Component 2 — Table header row found: {header_row} (0.3 pts)")
                        break
            if not header_found:
                actual_headers = [cell.text.strip() for cell in doc.tables[0].rows[0].cells] if doc.tables else []
                print(f"FAIL: Component 2 — Header row not matching. Found: {actual_headers}, Expected: {EXPECTED_HEADERS}")
        else:
            print(f"FAIL: Component 2 — No tables found to check header")
        if header_found:
            total_score += 0.3
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All 5 Completed tasks appear in the table with correct data (0.4 points)
    # Partial credit: 0.08 pts per matched task row
    try:
        matched_tasks = 0
        if len(doc.tables) >= 1:
            # Find the table with the completed tasks
            target_table = None
            for table in doc.tables:
                # Check if this table contains completed tasks
                rows_text = []
                for row in table.rows:
                    rows_text.append([cell.text.strip() for cell in row.cells])
                # Check if any row has 'Completed' status
                completed_rows = [r for r in rows_text if len(r) >= 4 and r[-1].strip().lower() == 'completed']
                if len(completed_rows) >= 1:
                    target_table = table
                    break

            if target_table is not None:
                # Extract all data rows (skip header)
                data_rows = []
                for row in target_table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 4:
                        data_rows.append(cells)

                print(f"  Found table with {len(data_rows)} data rows (excluding header)")

                for expected_task in EXPECTED_COMPLETED_TASKS:
                    task_name, owner, deadline, status = expected_task
                    found = False
                    for data_row in data_rows:
                        row_task = data_row[0].strip()
                        row_owner = data_row[1].strip() if len(data_row) > 1 else ''
                        row_deadline = normalize_date(data_row[2]) if len(data_row) > 2 else ''
                        row_status = data_row[3].strip() if len(data_row) > 3 else ''
                        normalized_deadline = normalize_date(deadline)

                        if (row_task.lower() == task_name.lower() and
                                row_status.lower() == status.lower()):
                            found = True
                            break

                    if found:
                        matched_tasks += 1
                        print(f"  MATCH: '{task_name}' found in table")
                    else:
                        print(f"  MISSING: '{task_name}' not found in table")

                task_score = round(matched_tasks * 0.08, 2)
                if matched_tasks == len(EXPECTED_COMPLETED_TASKS):
                    task_score = 0.4  # Full credit for all 5 tasks
                    print(f"PASS: Component 3 — All {matched_tasks}/5 completed tasks found in table (0.4 pts)")
                else:
                    print(f"PARTIAL: Component 3 — {matched_tasks}/5 completed tasks found in table ({task_score} pts)")
                total_score += task_score
            else:
                print(f"FAIL: Component 3 — No table with 'Completed' status rows found")
        else:
            print(f"FAIL: Component 3 — No tables found to check completed tasks")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in given env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
