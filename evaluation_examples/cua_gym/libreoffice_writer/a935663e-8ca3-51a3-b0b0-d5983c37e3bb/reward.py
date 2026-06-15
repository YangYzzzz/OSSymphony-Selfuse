"""
FINAL REWARD SCRIPT - SUCCESS
Task: All my top-level headings keep showing up in Liberation Serif, but the template requirements say every Heading 1 must be Times New Roman 16 pt Bold. How do I tweak the Heading 1 style so it instantly flips the whole document to that exact font, size, and weight?
Generated: 2025-09-10 12:32:17
Status: success
Model: azure-o3
Total Steps: 1
"""

from docx import Document
import os

# ---------------- Utility Helper Functions ----------------

def font_name_matches(name: str) -> bool:
    """Return True if the font name represents Times New Roman (case-insensitive)."""
    if not name:
        return False
    n = name.lower()
    return "times" in n and "roman" in n


def size_matches(size_obj) -> bool:
    """Return True if the size equals 16 pt (±0.5 tolerance)."""
    if size_obj is None:
        return False
    # python-docx returns length objects; grab .pt to get value in points
    try:
        pt_val = size_obj.pt
    except AttributeError:
        return False
    return abs(pt_val - 16) <= 0.5  # allow minimal rounding tolerance


def bold_matches(bold_val) -> bool:
    """Return True if bold is explicitly True (not None / False)."""
    return bool(bold_val)


# ---------------- Core Verification Function ----------------

def verify_heading1_style(doc_path: str) -> float:
    """Verify that:
    1. The Heading 1 style itself is Times New Roman 16 pt Bold (0.6 pts)
       • font name 0.2, size 0.2, bold 0.2
    2. Every paragraph that uses Heading 1 renders with those same effective
       properties (0.4 pts distributed proportionally).
    A progressive score (0.0-1.0) is returned.
    """
    print(f"Verifying Heading 1 style in document: {doc_path}")

    if not os.path.exists(doc_path):
        print("✗ File does not exist")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(doc_path)
        print(f"✓ Document loaded. Total paragraphs: {len(doc.paragraphs)}")
    except Exception as e:
        print(f"✗ Failed to load document: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- 1) Style Definition Check (max 0.6) ----------
    style_font_ok = style_size_ok = style_bold_ok = False
    try:
        h1_style = doc.styles["Heading 1"]
        font = h1_style.font
        style_font_ok = font_name_matches(font.name)
        style_size_ok = size_matches(font.size)
        style_bold_ok = bold_matches(font.bold)
        print("Heading 1 style attributes:")
        print(f"  Font name: {font.name} -> {'✓' if style_font_ok else '✗'}")
        print(f"  Font size: {font.size.pt if font.size else 'None'} pt -> {'✓' if style_size_ok else '✗'}")
        print(f"  Bold: {font.bold} -> {'✓' if style_bold_ok else '✗'}")
    except KeyError:
        print("✗ 'Heading 1' style not found in document")
    except Exception as exc:
        print(f"✗ Error accessing Heading 1 style: {exc}")

    score = 0.0
    if style_font_ok:
        score += 0.2
    if style_size_ok:
        score += 0.2
    if style_bold_ok:
        score += 0.2

    # ---------- 2) Paragraph-level Check (max 0.4) ----------
    h1_paras = [p for p in doc.paragraphs if p.style and p.style.name.lower().strip() == "heading 1"]
    print(f"Total Heading 1 paragraphs found: {len(h1_paras)}")

    if h1_paras:
        passed = 0
        for idx, para in enumerate(h1_paras, 1):
            para_ok = True
            for run in para.runs:
                # Effective (cascade) properties
                eff_name = run.font.name or (h1_style.font.name if 'h1_style' in locals() else None)
                eff_size = run.font.size or (h1_style.font.size if 'h1_style' in locals() else None)
                eff_bold = run.font.bold if run.font.bold is not None else (h1_style.font.bold if 'h1_style' in locals() else None)

                if not (font_name_matches(eff_name) and size_matches(eff_size) and bold_matches(eff_bold)):
                    para_ok = False
                    break
            print(f"  Paragraph {idx}: {'✓' if para_ok else '✗'} -> '{para.text[:50]}'")
            if para_ok:
                passed += 1
        ratio = passed / len(h1_paras)
        para_score = ratio * 0.4  # up to 0.4 points
        score += para_score
        print(f"Paragraph correctness ratio: {passed}/{len(h1_paras)} -> {para_score:.2f} points")
    else:
        print("✗ No Heading 1 paragraphs found; no points for paragraph checks.")

    final_score = min(score, 1.0)
    print(f"Total score: {final_score:.2f} (max 1.0)")
    print(f"REWARD: {final_score}")
    return final_score


# ---------------- Script Entrypoint ----------------
if __name__ == "__main__":
    verify_heading1_style("/home/user/all_my_top_level_headings_keep_showing_up_in_liberation_serif_but_the_template_requirements_say_ever.docx")
