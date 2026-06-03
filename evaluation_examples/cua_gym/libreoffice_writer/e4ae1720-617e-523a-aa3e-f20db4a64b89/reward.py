"""
Reward Script: Add MEAN formula in cell B12 for average of B2:B11, formatted to 2 decimal places
Task ID: writer_tm_044
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): B12 cell contains non-empty text with a numeric value
  Component 2 (0.35): The numeric value equals 80.40 (mean of the 10 scores)
  Component 3 (0.25): The displayed value shows exactly 2 decimal places
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_044'


def persist_app_state(domain):
    """Try to save any unsaved LibreOffice state."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print("PERSIST: ctrl+s sent for libreoffice_writer")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document has at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table has at least 12 rows and 2 columns
    if len(table.rows) < 12 or len(table.columns) < 2:
        print(f"FAIL: Table too small ({len(table.rows)} rows x {len(table.columns)} cols), need 12x2")
        print("REWARD: 0.0")
        return 0.0

    # Get B12 cell text (row index 11, col index 1)
    b12_text = table.cell(11, 1).text.strip()
    print(f"INFO: B12 cell text = {repr(b12_text)}")

    # Also check A12 to confirm it's the Average row (precondition gate, not scored)
    a12_text = table.cell(11, 0).text.strip()
    if 'average' not in a12_text.lower() and 'mean' not in a12_text.lower():
        print(f"WARN: A12 = {repr(a12_text)}, expected 'Average' label")

    # Component 1: B12 cell contains non-empty text with a numeric value (0.4 points)
    # This FAILS on initial (B12 is empty) and PASSES on golden (B12 = "80.40")
    try:
        numeric_value = None
        if b12_text:
            # Try to parse as a number (handle possible currency/formatting)
            cleaned = re.sub(r'[^\d.\-]', '', b12_text)
            if cleaned:
                numeric_value = float(cleaned)

        if numeric_value is not None:
            print(f"PASS: Component 1 -- B12 has numeric value {numeric_value} (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 -- B12 is empty or non-numeric: {repr(b12_text)}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: The numeric value equals 80.40 (mean of 72,88,65,91,84,77,93,69,86,79) (0.35 points)
    # Expected mean = (72+88+65+91+84+77+93+69+86+79) / 10 = 804 / 10 = 80.4
    try:
        if numeric_value is not None:
            expected_mean = 80.4
            if abs(numeric_value - expected_mean) < 0.05:
                print(f"PASS: Component 2 -- B12 value {numeric_value} matches expected mean {expected_mean} (0.35 pts)")
                total_score += 0.35
            else:
                print(f"FAIL: Component 2 -- B12 value {numeric_value} != expected mean {expected_mean}")
        else:
            print("FAIL: Component 2 -- No numeric value to check")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Display shows exactly 2 decimal places (0.25 points)
    # The task explicitly asks for 2 decimal places formatting.
    # "80.40" has 2 decimals, "80.4" has only 1.
    try:
        if b12_text:
            # Check for pattern: digits, decimal point, exactly 2 digits
            decimal_match = re.search(r'(\d+)\.(\d+)', b12_text)
            if decimal_match:
                decimal_part = decimal_match.group(2)
                if len(decimal_part) == 2:
                    print(f"PASS: Component 3 -- Value shows 2 decimal places: {repr(b12_text)} (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 3 -- Value has {len(decimal_part)} decimal places, expected 2: {repr(b12_text)}")
            else:
                print(f"FAIL: Component 3 -- No decimal point found in {repr(b12_text)}")
        else:
            print("FAIL: Component 3 -- B12 is empty")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
