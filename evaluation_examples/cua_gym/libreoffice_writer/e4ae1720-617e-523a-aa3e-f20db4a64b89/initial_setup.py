"""
Initial Setup: Add MEAN formula in table cell B12 for average of B2:B11
Task ID: writer_tm_044
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_044'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading('Mid-Term Exam Scores — Biology 201', level=1)

    # Introductory paragraph
    intro = doc.add_paragraph(
        'The following table lists the mid-term examination scores for '
        'Biology 201, Section B. Scores are out of 100 points.'
    )

    # Create 12-row x 2-column table
    table = doc.add_table(rows=12, cols=2)
    table.style = 'Table Grid'

    # Header row
    hdr_a = table.cell(0, 0)
    hdr_b = table.cell(0, 1)
    hdr_a.text = ''
    hdr_b.text = ''
    run_a = hdr_a.paragraphs[0].add_run('Student')
    run_a.bold = True
    run_a.font.size = Pt(11)
    run_b = hdr_b.paragraphs[0].add_run('Score')
    run_b.bold = True
    run_b.font.size = Pt(11)

    # Student data — rows 2 through 11 (0-indexed: rows 1-10)
    students = [
        ('Emily Rodriguez', 72),
        ('James Nakamura', 88),
        ('Priya Sharma', 65),
        ('Liam O\'Brien', 91),
        ('Sofia Andersson', 84),
        ('Wei Zhang', 77),
        ('Amara Okafor', 93),
        ('Daniel Petrov', 69),
        ('Hannah Kim', 86),
        ('Lucas Fernandez', 79),
    ]

    for i, (name, score) in enumerate(students):
        row_idx = i + 1  # rows 1..10
        table.cell(row_idx, 0).text = name
        table.cell(row_idx, 1).text = str(score)

    # Row 12 (index 11): "Average" label, B12 left EMPTY
    table.cell(11, 0).text = 'Average'
    # B12 is intentionally left empty — the agent must add the MEAN formula

    # Closing paragraph
    doc.add_paragraph(
        'Please calculate the class average and enter it in the table above.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
