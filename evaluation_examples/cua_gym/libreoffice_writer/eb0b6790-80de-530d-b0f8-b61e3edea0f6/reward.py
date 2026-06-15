"""
FINAL REWARD SCRIPT - SUCCESS
Task: Paragraph 4 is a mess—some sentences have double spaces after the period, others have none. In LibreOffice Writer, how can I run a quick find-and-replace (with regular expressions enabled) that scans only that paragraph and forces every full stop to be followed by exactly one space?
Generated: 2025-09-10 18:19:43
Status: success
Model: azure-o3
Total Steps: 5
"""

import os
import re
from docx import Document


def verify_single_space_after_period(doc_path: str) -> float:
    """Verify that, in the target document, paragraph 4 (0-based index 4)
    has exactly one space following every full-stop (period) that is not the
    last character of the paragraph.

    Scoring (progressive):
        0.7  – every period followed by exactly one space (or EOL)
        0.3  – paragraph contains no double-spaces at all
        1.0  – both conditions satisfied
    Returns a float between 0.0 and 1.0 and prints a detailed breakdown.
    """

    print(f"Verifying document: {doc_path}")

    # ---------- prerequisite: file must exist & load ----------
    if not os.path.exists(doc_path):
        print("✗ File not found – cannot verify")
        return 0.0  # no score if file absent

    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"✗ Error loading DOCX: {e}")
        return 0.0

    # ---------- locate the specific paragraph (paragraph 4) ----------
    if len(doc.paragraphs) < 5:
        print("✗ Document has fewer than 5 paragraphs – target paragraph missing")
        return 0.0

    target_para = doc.paragraphs[4].text
    print("Target paragraph text:")
    print(target_para)

    total_score = 0.0  # progressive scoring accumulator

    # ---------- Requirement 1: every period followed by exactly 1 space ----------
    period_positions = [m.start() for m in re.finditer(r'\.', target_para)]
    if not period_positions:
        print("✗ No periods found – cannot assess spacing")
    else:
        incorrect = 0
        correct = 0
        for pos in period_positions:
            if pos == len(target_para) - 1:
                # Period at end of paragraph – automatically correct
                correct += 1
            else:
                next_char = target_para[pos + 1]
                if next_char == ' ':
                    # Ensure there is NOT a second consecutive space
                    if (pos + 2) < len(target_para) and target_para[pos + 2] == ' ':
                        incorrect += 1  # double-space detected
                    else:
                        correct += 1
                else:
                    incorrect += 1  # no space after period

        print(f"Total periods: {len(period_positions)}; Correctly spaced: {correct}; Incorrect: {incorrect}")
        if incorrect == 0:
            total_score += 0.7
            print("✓ All periods correctly followed by one space (0.7 points)")
        else:
            print("✗ Some periods incorrectly spaced (0 points for this criterion)")

    # ---------- Requirement 2: paragraph contains NO double spaces anywhere ----------
    if '  ' not in target_para:
        total_score += 0.3
        print("✓ No double spaces found in paragraph (0.3 points)")
    else:
        print("✗ Double spaces still present (0 points for this criterion)")

    final_score = min(total_score, 1.0)
    print(f"Total score: {final_score}")
    return final_score


if __name__ == "__main__":
    FILE_PATH = "/home/user/paragraph_4_is_a_messsome_sentences_have_double_spaces_after_the_period_others_have_none_in_libreoff.docx"
    reward = verify_single_space_after_period(FILE_PATH)
    print(f"REWARD: {reward}")
