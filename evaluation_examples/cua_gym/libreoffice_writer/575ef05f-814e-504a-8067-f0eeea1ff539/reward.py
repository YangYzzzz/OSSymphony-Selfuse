"""
Reward Script: School newsletter in LibreOffice Writer
Task ID: writer_wf_024
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Title "Oakridge Elementary - Parent Newsletter" centered
  Component 2 (0.20): Four section headings present (Principal's Message, Upcoming Events, Student Achievement, Important Dates)
  Component 3 (0.20): 5 bulleted list items under Upcoming Events
  Component 4 (0.20): Important Dates table with Date/Event columns and 4 data rows
  Component 5 (0.10): Footer with school contact information
  Component 6 (0.10): 2-column layout in the document
"""

import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_024'


def persist_app_state(domain):
    """Best-effort save via Ctrl+S in case document is still open in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gather all paragraph texts for reuse
    all_para_texts = [p.text.strip() for p in doc.paragraphs]

    # Component 1: Title "Oakridge Elementary - Parent Newsletter" centered (0.20 points)
    try:
        title_found = False
        for p in doc.paragraphs:
            text_lower = p.text.strip().lower()
            if 'oakridge elementary' in text_lower and 'newsletter' in text_lower:
                # Check centered alignment
                align = p.paragraph_format.alignment
                if align == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    print(f"PASS: Component 1 — Title found centered: '{p.text.strip()[:60]}' (0.20 pts)")
                    total_score += 0.20
                    title_found = True
                    break
                else:
                    print(f"PARTIAL: Component 1 — Title found but not centered (align={align})")
                    total_score += 0.10
                    title_found = True
                    break
        if not title_found:
            print("FAIL: Component 1 — No paragraph with 'Oakridge Elementary' and 'newsletter'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Four section headings present (0.20 points)
    # Each heading is worth 0.05
    try:
        required_headings = [
            "principal's message",
            "upcoming events",
            "student achievement",
            "important dates",
        ]
        heading_score = 0.0
        found_headings = []
        for p in doc.paragraphs:
            txt = p.text.strip().lower()
            style_name = p.style.name.lower() if p.style else ''
            for rh in required_headings:
                if rh in txt and rh not in found_headings:
                    # Accept heading style or bold formatting as section heading
                    is_heading_style = 'heading' in style_name
                    is_bold = any(r.font.bold for r in p.runs if r.text.strip())
                    if is_heading_style or is_bold or len(txt) < 40:
                        found_headings.append(rh)
                        heading_score += 0.05

        if heading_score > 0:
            print(f"PASS: Component 2 — Found {len(found_headings)}/4 headings: {found_headings} ({heading_score:.2f} pts)")
        else:
            print(f"FAIL: Component 2 — No required headings found")
        total_score += heading_score
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 5 bulleted list items under Upcoming Events (0.20 points)
    try:
        bullet_count = 0
        in_events_section = False
        for p in doc.paragraphs:
            txt = p.text.strip().lower()
            # Track when we enter Upcoming Events section
            if 'upcoming events' in txt:
                in_events_section = True
                continue
            # Track when we leave (next heading)
            if in_events_section and p.style and 'heading' in p.style.name.lower():
                break
            # Count bullets in the Upcoming Events section
            if in_events_section:
                style_name = p.style.name.lower() if p.style else ''
                if 'list' in style_name or 'bullet' in style_name:
                    bullet_count += 1

        if bullet_count >= 5:
            print(f"PASS: Component 3 — Found {bullet_count} bullet items in Upcoming Events (0.20 pts)")
            total_score += 0.20
        elif bullet_count >= 3:
            partial = round(bullet_count * 0.04, 2)
            print(f"PARTIAL: Component 3 — Found {bullet_count}/5 bullets ({partial} pts)")
            total_score += partial
        elif bullet_count > 0:
            partial = round(bullet_count * 0.04, 2)
            print(f"PARTIAL: Component 3 — Found {bullet_count}/5 bullets ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No bullet items found in Upcoming Events section")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Important Dates table with Date/Event columns and 4 data rows (0.20 points)
    try:
        table_score = 0.0
        if len(doc.tables) >= 1:
            # Find the table that has Date/Event headers
            target_table = None
            for t in doc.tables:
                if len(t.columns) >= 2:
                    header_texts = [c.text.strip().lower() for c in t.rows[0].cells]
                    if 'date' in header_texts and 'event' in header_texts:
                        target_table = t
                        break

            if target_table is None:
                # Accept any table with 2 columns as a fallback
                for t in doc.tables:
                    if len(t.columns) == 2:
                        target_table = t
                        break

            if target_table is not None:
                num_rows = len(target_table.rows)
                num_cols = len(target_table.columns)
                # Check structure: 2 columns
                if num_cols >= 2:
                    table_score += 0.05
                    print(f"  Table has {num_cols} columns (expected 2): OK")

                # Check header row has Date and Event
                header_texts = [c.text.strip().lower() for c in target_table.rows[0].cells]
                if 'date' in header_texts and 'event' in header_texts:
                    table_score += 0.05
                    print(f"  Header row: {header_texts}: OK")
                else:
                    print(f"  Header row: {header_texts}: missing Date/Event")

                # Check data rows (should be 4, total 5 rows including header)
                data_rows = num_rows - 1  # subtract header
                if data_rows >= 4:
                    table_score += 0.10
                    print(f"  Data rows: {data_rows} (expected 4): OK")
                elif data_rows >= 2:
                    table_score += 0.05
                    print(f"  Data rows: {data_rows} (expected 4): partial")
                else:
                    print(f"  Data rows: {data_rows} (expected 4): insufficient")
            else:
                print("  No suitable table found")

            if table_score > 0:
                print(f"PASS: Component 4 — Table verified ({table_score:.2f} pts)")
            else:
                print(f"FAIL: Component 4 — Table structure issues")
        else:
            print("FAIL: Component 4 — No tables found in document")
        total_score += table_score
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Footer with school contact information (0.10 points)
    try:
        footer_found = False
        for sec in doc.sections:
            ft = sec.footer
            if ft.is_linked_to_previous and sec != doc.sections[0]:
                continue
            for fp in ft.paragraphs:
                footer_text = fp.text.strip().lower()
                if footer_text and ('oakridge' in footer_text or 'phone' in footer_text or 'email' in footer_text or 'school' in footer_text):
                    footer_found = True
                    break
            if footer_found:
                break

        if footer_found:
            print(f"PASS: Component 5 — Footer with contact info found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 5 — No footer with contact information found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: 2-column layout (0.10 points)
    try:
        two_col_found = False
        for sec in doc.sections:
            cols_elems = sec._sectPr.findall(qn('w:cols'))
            for c in cols_elems:
                num = c.get(qn('w:num'))
                if num is not None and int(num) >= 2:
                    two_col_found = True
                    break
            if two_col_found:
                break

        if two_col_found:
            print(f"PASS: Component 6 — 2-column layout found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 6 — No 2-column layout detected")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
