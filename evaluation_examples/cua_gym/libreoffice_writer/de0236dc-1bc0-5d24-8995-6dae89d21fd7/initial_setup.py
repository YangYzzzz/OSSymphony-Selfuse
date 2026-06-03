"""
Initial Setup: Economics textbook document with key terms throughout
Task ID: writer_mt_079
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_079'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)

# --- Economics textbook chapter content ---
# Each chapter contains multiple key terms naturally woven in.

CHAPTERS = [
    {
        "title": "Chapter 1: Introduction to Economics",
        "sections": [
            {
                "heading": "What is Economics?",
                "paragraphs": [
                    "Economics is the social science that studies how individuals, businesses, governments, and societies allocate scarce resources. At its core, economics examines the forces of supply and demand that drive market behavior. Understanding these fundamental concepts helps us make sense of the complex global marketplace.",
                    "The study of economics is divided into two major branches: microeconomics and macroeconomics. Microeconomics focuses on individual decision-making, including how consumers respond to changes in supply and demand for specific goods and services. Macroeconomics takes a broader view, examining national output measured by GDP, unemployment rates, and inflation across an entire economy.",
                    "One of the most important concepts in economics is equilibrium, the state where supply equals demand at a given price level. When a market reaches equilibrium, there is no tendency for the price to change unless an external factor shifts either the supply curve or the demand curve. This concept of equilibrium underpins much of economic analysis.",
                ]
            },
            {
                "heading": "The Role of Markets",
                "paragraphs": [
                    "Markets are mechanisms through which buyers and sellers interact to exchange goods and services. In a perfectly competitive market, supply and demand interact freely, and prices adjust until equilibrium is reached. However, perfect competition is rare in the real world.",
                    "Market structures vary widely. A monopoly exists when a single firm dominates an entire industry, controlling supply and setting prices above competitive levels. An oligopoly occurs when a small number of large firms dominate a market, often engaging in strategic behavior. Both monopoly and oligopoly represent departures from the ideal of perfect competition.",
                    "Government intervention in markets can take many forms. A tariff is a tax imposed on imported goods, designed to protect domestic industries from foreign competition. A subsidy is a payment made by the government to producers, intended to lower production costs and encourage output. Both tariff and subsidy policies affect the equilibrium price and quantity in affected markets.",
                ]
            },
        ]
    },
    {
        "title": "Chapter 2: Measuring Economic Performance",
        "sections": [
            {
                "heading": "Gross Domestic Product",
                "paragraphs": [
                    "Gross Domestic Product, commonly known as GDP, is the total market value of all final goods and services produced within a country's borders during a specific time period, typically one year. GDP serves as the primary indicator of a nation's economic health and is closely watched by policymakers, investors, and analysts worldwide.",
                    "GDP can be measured using three approaches: the expenditure approach, the income approach, and the production approach. Each method should yield the same result. Real GDP adjusts for inflation, allowing meaningful comparisons across time periods. When real GDP grows, the economy is expanding; when it contracts, the economy may be entering a recession.",
                    "Economists track GDP growth rates to identify economic cycles. A sustained decline in GDP over two consecutive quarters is the traditional definition of a recession. A more severe and prolonged downturn is classified as a depression. The Great Depression of the 1930s saw GDP fall by nearly 30%, illustrating the devastating impact of a full-blown depression on society.",
                ]
            },
            {
                "heading": "Inflation and Deflation",
                "paragraphs": [
                    "Inflation refers to a general increase in the price level of goods and services over time. When inflation rises, each unit of currency buys fewer goods and services. Central banks aim to maintain inflation at a moderate level, typically around 2% per year in developed economies.",
                    "The Consumer Price Index (CPI) is the most widely used measure of inflation. It tracks the cost of a fixed basket of goods and services over time. When the CPI rises, inflation is occurring. Hyperinflation, an extreme form of inflation, can destroy an economy's monetary system entirely.",
                    "Deflation is the opposite of inflation: a general decrease in the price level. While lower prices might seem beneficial to consumers, persistent deflation can be harmful to an economy. Deflation increases the real burden of debt and can lead to a deflationary spiral where falling prices reduce business revenue, leading to layoffs and further reductions in demand. Japan experienced prolonged deflation during its 'Lost Decade' of the 1990s.",
                    "The relationship between inflation, deflation, and monetary policy is central to macroeconomic management. Central banks use the interest rate as their primary tool to influence inflation and deflation. Raising the interest rate tends to reduce inflation by making borrowing more expensive, while lowering the interest rate can combat deflation by encouraging spending.",
                ]
            },
        ]
    },
    {
        "title": "Chapter 3: Microeconomic Foundations",
        "sections": [
            {
                "heading": "Supply, Demand, and Elasticity",
                "paragraphs": [
                    "The law of demand states that, all else being equal, as the price of a good increases, the quantity demanded decreases. Conversely, as the price falls, demand rises. This inverse relationship between price and demand is one of the most fundamental principles in economics.",
                    "The law of supply states the opposite: as the price of a good increases, the quantity supplied also increases. Producers are willing to supply more at higher prices because the potential for profit is greater. The interaction of supply and demand determines the market equilibrium price.",
                    "Elasticity measures the responsiveness of supply or demand to changes in price or income. Price elasticity of demand measures how much the quantity demanded changes when the price changes. If a small price increase causes a large drop in demand, the good has high elasticity. Necessities like food tend to have low elasticity, while luxury goods have high elasticity.",
                    "Cross-price elasticity measures how the demand for one good responds to a change in the price of another good. Income elasticity measures how demand changes with consumer income. Understanding elasticity helps businesses set pricing strategies and helps governments predict the impact of taxes and subsidies on market outcomes.",
                ]
            },
            {
                "heading": "Market Structures",
                "paragraphs": [
                    "Perfect competition represents one extreme of market structure. In a perfectly competitive market, many small firms sell identical products, and no single firm can influence the market price. Agricultural markets sometimes approximate perfect competition.",
                    "At the other extreme, a monopoly exists when a single firm is the sole producer of a product with no close substitutes. A monopoly firm is a price maker rather than a price taker. Natural monopoly arises when economies of scale make it most efficient for a single firm to serve the entire market, as in utilities. Regulation of monopoly is a key function of government.",
                    "An oligopoly is a market structure dominated by a few large firms. In an oligopoly, each firm's decisions affect the others, leading to strategic interdependence. Oligopoly firms may engage in collusion, forming cartels to set prices and output levels, or they may compete aggressively. The airline and telecommunications industries are common examples of oligopoly.",
                    "Monopolistic competition lies between perfect competition and oligopoly. Many firms sell differentiated products, giving each some degree of market power. Restaurants, clothing brands, and personal care products operate in monopolistically competitive markets.",
                ]
            },
        ]
    },
    {
        "title": "Chapter 4: Macroeconomic Policy",
        "sections": [
            {
                "heading": "Fiscal Policy",
                "paragraphs": [
                    "Fiscal policy refers to the government's use of taxation and spending to influence the economy. Expansionary fiscal policy involves increasing government spending or reducing taxes to stimulate economic growth during a recession. Contractionary fiscal policy involves decreasing spending or raising taxes to cool an overheating economy and reduce inflation.",
                    "The effectiveness of fiscal policy depends on several factors, including the size of the government spending multiplier, the current state of the economy, and the response of monetary policy. During the 2008 financial crisis, many governments implemented large fiscal policy stimulus packages to prevent a deeper recession or depression.",
                    "Government budget deficits occur when spending exceeds revenue. Persistent deficits lead to growing national debt. Critics of expansionary fiscal policy argue that excessive debt can crowd out private investment and burden future generations. Proponents counter that fiscal policy is essential for stabilizing the economy during downturns when monetary policy alone may be insufficient.",
                ]
            },
            {
                "heading": "Monetary Policy",
                "paragraphs": [
                    "Monetary policy is conducted by central banks and involves managing the money supply and the interest rate to achieve macroeconomic objectives such as price stability, full employment, and economic growth. The Federal Reserve in the United States, the European Central Bank, and the Bank of Japan are among the world's most influential monetary policy institutions.",
                    "The primary tool of monetary policy is the setting of short-term interest rate targets. When a central bank lowers the interest rate, borrowing becomes cheaper, encouraging businesses to invest and consumers to spend. This expansionary monetary policy helps combat recession and deflation. Conversely, raising the interest rate makes borrowing more expensive, slowing economic activity and helping to reduce inflation.",
                    "Unconventional monetary policy tools include quantitative easing, forward guidance, and negative interest rate policies. During the 2008 crisis and the COVID-19 pandemic, central banks deployed these tools when traditional monetary policy reached the zero lower bound on the interest rate. The effectiveness of these unconventional approaches remains a subject of debate among economists.",
                ]
            },
        ]
    },
    {
        "title": "Chapter 5: International Economics",
        "sections": [
            {
                "heading": "Trade and Tariffs",
                "paragraphs": [
                    "International trade allows countries to specialize in producing goods where they have a comparative advantage. By trading with other nations, countries can consume beyond their production possibilities. The theory of comparative advantage, first articulated by David Ricardo, remains the foundation of international trade theory.",
                    "A tariff is a tax on imported goods. Governments impose tariff barriers for several reasons: to protect domestic industries from foreign competition, to raise revenue, or to retaliate against trading partners. While a tariff benefits domestic producers of the protected good, it raises prices for consumers and reduces overall economic efficiency.",
                    "A trade deficit occurs when a country imports more goods and services than it exports. The United States has run a persistent trade deficit for decades, importing far more than it exports. While a trade deficit is often viewed negatively in political discourse, economists debate whether a trade deficit is truly harmful. Some argue that a trade deficit reflects strong consumer demand and capital inflows, while others worry about the accumulation of foreign debt.",
                    "The exchange rate is the price of one currency in terms of another. Changes in the exchange rate affect international trade flows. When a country's currency depreciates (its exchange rate falls), its exports become cheaper and imports become more expensive, tending to reduce the trade deficit. A strong exchange rate has the opposite effect, making imports cheaper but exports less competitive.",
                ]
            },
            {
                "heading": "Subsidies and Trade Policy",
                "paragraphs": [
                    "A subsidy is a financial contribution from the government to domestic producers. Export subsidy programs help domestic firms compete in international markets by lowering their costs. Agricultural subsidy programs are common in both developed and developing countries, though they distort global trade patterns.",
                    "The World Trade Organization (WTO) sets rules for international trade, including restrictions on tariff levels and subsidy programs. Trade agreements aim to reduce tariff barriers and create a more level playing field. However, disputes over tariff and subsidy policies remain a frequent source of tension between trading partners.",
                    "Regional trade agreements, such as the USMCA and the European Union's single market, eliminate tariff barriers among member countries. These agreements promote economic integration but can create trade diversion, where imports shift from efficient non-member producers to less efficient member producers who benefit from the tariff-free access.",
                ]
            },
        ]
    },
    {
        "title": "Chapter 6: Financial Markets and Instruments",
        "sections": [
            {
                "heading": "Interest Rates and Investment",
                "paragraphs": [
                    "The interest rate is the cost of borrowing money, typically expressed as an annual percentage of the principal. Interest rate levels are determined by the interaction of the supply of and demand for loanable funds, along with central bank monetary policy decisions. The interest rate plays a crucial role in investment decisions by both businesses and households.",
                    "When the interest rate is low, businesses find it cheaper to finance new projects, and consumers are more likely to take out mortgages and auto loans. Low interest rate environments tend to boost economic activity but can also lead to asset bubbles and excessive risk-taking. The relationship between the interest rate and investment spending is a key channel through which monetary policy affects GDP.",
                    "Bond prices and the interest rate move inversely. When the interest rate rises, existing bonds with lower coupon rates become less attractive, and their prices fall. This inverse relationship between the interest rate and bond prices is fundamental to understanding financial markets.",
                ]
            },
            {
                "heading": "Capital Markets",
                "paragraphs": [
                    "Capital gains refer to the profit earned when an asset is sold for more than its purchase price. Investors in stocks, real estate, and other assets seek capital gains as a primary source of return. The taxation of capital gains varies by jurisdiction and holding period. Long-term capital gains, from assets held over one year, are often taxed at lower rates than short-term capital gains.",
                    "Capital gains taxes influence investor behavior and portfolio allocation. When capital gains tax rates are high, investors may hold assets longer to defer taxation, reducing market liquidity. Conversely, lower capital gains tax rates encourage more frequent trading and investment. The optimal capital gains tax rate is a perennial debate in fiscal policy.",
                    "Amortization is the process of spreading the cost of an intangible asset or a loan over its useful life or loan term. In accounting, amortization applies to intangible assets such as patents, trademarks, and goodwill. In finance, amortization refers to the gradual repayment of a loan through regular installments that cover both principal and the interest rate.",
                    "An amortization schedule shows how each payment is divided between principal and interest over the life of a loan. In the early years, a larger portion of each payment goes toward interest, while in later years, more goes toward principal. Understanding amortization is essential for both corporate financial planning and personal mortgage management. The concept of amortization also appears in the calculation of GDP when accounting for the depreciation of capital assets.",
                ]
            },
        ]
    },
    {
        "title": "Chapter 7: Economic Downturns",
        "sections": [
            {
                "heading": "Recessions",
                "paragraphs": [
                    "A recession is a significant decline in economic activity that lasts for an extended period, typically defined as two or more consecutive quarters of declining GDP. During a recession, unemployment rises, consumer spending falls, and business investment contracts. Recession can be triggered by various factors, including financial crises, supply shocks, or abrupt changes in monetary policy.",
                    "The National Bureau of Economic Research (NBER) officially dates recessions in the United States. A recession typically involves declines in GDP, employment, industrial production, and retail sales. The 2008 recession, triggered by the collapse of the housing market and the subprime mortgage crisis, was the most severe recession since the Great Depression.",
                    "Policy responses to recession typically include expansionary fiscal policy (increased government spending, tax cuts) and expansionary monetary policy (lower interest rate targets, quantitative easing). The speed and effectiveness of these responses can determine whether a recession is brief and mild or prolonged and severe.",
                ]
            },
            {
                "heading": "Depressions",
                "paragraphs": [
                    "A depression is an extreme and prolonged economic downturn that goes far beyond a typical recession. While there is no universally agreed-upon definition, a depression is generally characterized by a decline in GDP of more than 10%, unemployment rates exceeding 20%, and a duration of several years. The Great Depression of the 1930s remains the defining example.",
                    "During the Great Depression, GDP in the United States fell by approximately 30% from its peak. Deflation was rampant, with prices falling by more than 25%. The depression spread globally through trade and financial linkages. Bank failures wiped out savings, and unemployment reached 25%. The human suffering caused by the depression led to fundamental changes in economic thinking and government policy.",
                    "Keynesian economics emerged in response to the Great Depression. John Maynard Keynes argued that during a depression, private demand is insufficient to maintain full employment, and government fiscal policy must fill the gap. His ideas revolutionized macroeconomic policy and remain influential today. Modern economists generally believe that the aggressive use of fiscal policy and monetary policy can prevent a recession from deepening into a full depression.",
                ]
            },
        ]
    },
    {
        "title": "Chapter 8: Price Dynamics and Exchange",
        "sections": [
            {
                "heading": "Exchange Rates and Global Commerce",
                "paragraphs": [
                    "The exchange rate between two currencies determines how much of one currency is needed to purchase a unit of another. Exchange rate movements have far-reaching effects on international trade, investment flows, and economic policy. A country's exchange rate regime can be fixed, floating, or managed, each with distinct advantages and risks.",
                    "Under a floating exchange rate system, currency values are determined by market forces of supply and demand. When demand for a currency increases relative to its supply, the exchange rate appreciates. Factors that influence exchange rate movements include interest rate differentials, inflation rates, GDP growth, political stability, and market sentiment.",
                    "Purchasing power parity (PPP) is a theory that suggests exchange rate levels should adjust so that identical goods cost the same in different countries. While PPP holds approximately in the long run, short-term exchange rate movements often deviate significantly from PPP predictions due to capital flows, speculation, and differences in interest rate policies across countries.",
                ]
            },
            {
                "heading": "Trade Balances",
                "paragraphs": [
                    "A country's trade balance is the difference between the value of its exports and imports. A trade deficit occurs when imports exceed exports, while a trade surplus occurs when exports exceed imports. The trade deficit of the United States has been a persistent feature of the global economy, driven by strong consumer demand for imported goods.",
                    "The trade deficit is influenced by several factors, including the exchange rate, relative GDP growth rates, trade agreements, tariff levels, and subsidy policies. A weaker exchange rate tends to improve the trade deficit by making exports cheaper and imports more expensive, though the adjustment process can take time due to contractual lags.",
                    "Capital account flows offset the trade deficit. When a country runs a trade deficit, it must be financing that deficit through capital inflows, such as foreign purchases of domestic bonds and equities. This connection between the trade deficit and capital flows is captured by the balance of payments identity, a fundamental relationship in international economics.",
                ]
            },
        ]
    },
]

def create_initial():
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title page
    title_para = doc.add_heading("Economics 101: Principles and Applications", level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("A Comprehensive Introduction to Economic Theory")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run("Dr. Alexandra Mitchell\nProfessor of Economics\nStanford University")
    run.font.size = Pt(12)

    edition = doc.add_paragraph()
    edition.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = edition.add_run("Fourth Edition, 2025")
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_page_break()

    # Table of Contents placeholder
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for i, chapter in enumerate(CHAPTERS, 1):
        toc_entry = doc.add_paragraph()
        run = toc_entry.add_run(chapter["title"])
        run.font.size = Pt(11)
        for section_data in chapter["sections"]:
            sub_entry = doc.add_paragraph()
            sub_entry.paragraph_format.left_indent = Inches(0.5)
            run = sub_entry.add_run(section_data["heading"])
            run.font.size = Pt(10)

    doc.add_page_break()

    # Preface
    doc.add_heading("Preface", level=1)
    preface_text = (
        "This textbook provides a comprehensive introduction to the principles of economics. "
        "It covers both microeconomic and macroeconomic topics, including supply and demand analysis, "
        "market structures such as monopoly and oligopoly, fiscal policy and monetary policy, "
        "international trade concepts including tariff and subsidy effects, and financial instruments "
        "such as capital gains and amortization. The text emphasizes real-world applications and "
        "data-driven analysis to help students understand how economic forces shape our daily lives."
    )
    doc.add_paragraph(preface_text)

    preface2 = (
        "Key topics covered include GDP measurement, inflation and deflation dynamics, "
        "the causes and consequences of recession and depression, exchange rate determination, "
        "trade deficit analysis, interest rate policy, elasticity of demand and supply, "
        "and the equilibrium conditions that govern market outcomes. Each chapter builds on "
        "the previous one, creating a cohesive framework for understanding modern economics."
    )
    doc.add_paragraph(preface2)

    doc.add_page_break()

    # Chapters
    for chapter in CHAPTERS:
        doc.add_heading(chapter["title"], level=1)
        for section_data in chapter["sections"]:
            doc.add_heading(section_data["heading"], level=2)
            for para_text in section_data["paragraphs"]:
                doc.add_paragraph(para_text)
        doc.add_page_break()

    # Glossary (short, no index)
    doc.add_heading("Glossary of Key Terms", level=1)
    glossary_terms = [
        ("Amortization", "The process of spreading the cost of an intangible asset or loan over its useful life or term."),
        ("Capital Gains", "The profit realized from the sale of an asset at a price higher than its purchase price."),
        ("Deflation", "A general decline in the price level of goods and services."),
        ("Demand", "The quantity of a good or service that consumers are willing and able to purchase at various prices."),
        ("Depression", "A severe and prolonged economic downturn, more extreme than a recession."),
        ("Elasticity", "A measure of the responsiveness of quantity demanded or supplied to a change in price or income."),
        ("Equilibrium", "The state in which market supply and demand balance each other, resulting in stable prices."),
        ("Exchange Rate", "The price of one currency expressed in terms of another currency."),
        ("Fiscal Policy", "Government use of taxation and spending to influence the economy."),
        ("GDP (Gross Domestic Product)", "The total market value of all final goods and services produced within a country in a given period."),
        ("Inflation", "A general increase in the price level of goods and services over time."),
        ("Interest Rate", "The cost of borrowing money, expressed as an annual percentage of the principal."),
        ("Monetary Policy", "Central bank actions to manage the money supply and interest rates."),
        ("Monopoly", "A market structure in which a single firm is the sole producer of a product with no close substitutes."),
        ("Oligopoly", "A market structure dominated by a small number of large firms."),
        ("Recession", "A significant decline in economic activity lasting more than a few months."),
        ("Subsidy", "A government payment to producers to lower production costs and encourage output."),
        ("Supply", "The quantity of a good or service that producers are willing and able to offer at various prices."),
        ("Tariff", "A tax imposed on imported goods."),
        ("Trade Deficit", "The amount by which a country's imports exceed its exports."),
    ]
    for term, definition in glossary_terms:
        p = doc.add_paragraph()
        run_term = p.add_run(f"{term}: ")
        run_term.bold = True
        run_term.font.size = Pt(11)
        run_def = p.add_run(definition)
        run_def.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')

create_initial()
