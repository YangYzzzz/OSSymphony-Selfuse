"""
Initial Setup: Employee handbook section — strikethrough task pre-state
Task ID: osworld_writer_strikethrough_last_para_006
Domain: libreoffice_writer

Creates a 4-paragraph employee handbook section. The last paragraph
describes an outdated policy and must NOT have strikethrough applied
(the agent's task is to apply it).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Section heading
    heading = doc.add_heading("Employee Handbook — Section 4: Leave and Time-Off Policies", level=1)

    # Paragraph 1: Current policy introduction
    p1 = doc.add_paragraph(
        "All full-time employees are entitled to a minimum of fifteen (15) days of paid annual "
        "leave per calendar year, accruing at a rate of 1.25 days per month of continuous service. "
        "Leave requests must be submitted through the HR portal at least five (5) business days in "
        "advance and are subject to manager approval based on operational requirements."
    )

    # Paragraph 2: Sick leave policy
    p2 = doc.add_paragraph(
        "Employees may take up to ten (10) days of paid sick leave annually without requiring a "
        "medical certificate. For absences exceeding three (3) consecutive days, a certificate from "
        "a licensed medical practitioner must be submitted to Human Resources within five (5) "
        "business days of returning to work. Unused sick leave does not carry over to the following year."
    )

    # Paragraph 3: Parental leave
    p3 = doc.add_paragraph(
        "Primary caregivers are eligible for up to sixteen (16) weeks of paid parental leave "
        "following the birth, adoption, or foster placement of a child. Secondary caregivers are "
        "entitled to four (4) weeks of paid parental leave. Employees must notify HR at least eight "
        "(8) weeks prior to the anticipated leave start date and provide relevant documentation."
    )

    # Paragraph 4: OUTDATED policy — must NOT have strikethrough in initial state
    p4 = doc.add_paragraph(
        "Employees who have completed at least two (2) years of continuous service are eligible "
        "to apply for a one-time sabbatical leave of up to three (3) months at fifty percent (50%) "
        "of their base salary. Sabbatical requests are approved solely at the discretion of the "
        "department head and must be submitted no later than six (6) months in advance. This policy "
        "is superseded by the updated Flexible Work Arrangement guidelines effective January 2024."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
