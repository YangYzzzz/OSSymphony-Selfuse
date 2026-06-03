"""
Reward Script: Mark final paragraph with strikethrough formatting
Task ID: osworld_writer_strikethrough_last_para_006
Domain: libreoffice_writer
Scoring:
  Component 1 (0.7): All runs with text in the last paragraph have strike=True
  Component 2 (0.3): The last paragraph has strikethrough AND its text content is
                      intact (not truncated or corrupted by the formatting action)
                      — compound check anchored to the task change
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_strikethrough_last_para_006'

# Expected last paragraph content (from task context: the outdated sabbatical policy)
# Used in Component 2 to verify text was not corrupted during formatting
EXPECTED_LAST_PARA_FRAGMENTS = [
    "Employees who have completed at least two (2) years",
    "sabbatical leave",
    "January 2024",
]


def persist_app_state():
    """Send Ctrl+S to save any unsaved LibreOffice edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Apply strikethrough formatting to every character in the last paragraph
    of the employee handbook document (4 content paragraphs + 1 heading = 5 total).
    """
    total_score = 0.0

    # Precondition gate: can we load the file?
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: document must have at least 2 paragraphs (heading + at least 1 content)
    if len(doc.paragraphs) < 2:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraph(s); expected at least 2.")
        print("REWARD: 0.0")
        return 0.0

    last_para = doc.paragraphs[-1]

    # Component 1: All runs with text in the last paragraph have strike=True (0.7 points)
    # This is the primary task requirement: every word/character in the final paragraph
    # must have strikethrough formatting applied.
    # FAILS on initial_env (strike=None) → PASSES on golden_env (strike=True)
    try:
        runs_with_text = [run for run in last_para.runs if run.text.strip()]
        if len(runs_with_text) == 0:
            print("FAIL: Component 1 — Last paragraph has no runs with text; cannot verify strikethrough.")
        else:
            all_strike = all(run.font.strike is True for run in runs_with_text)
            if all_strike:
                print(f"PASS: Component 1 — All {len(runs_with_text)} run(s) in last paragraph have strike=True (0.7 pts)")
                total_score += 0.7
            else:
                non_strike = [repr(r.text[:40]) for r in runs_with_text if r.font.strike is not True]
                print(f"FAIL: Component 1 — Some runs in last paragraph do NOT have strike=True. Non-strikethrough: {non_strike}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: The last paragraph has strikethrough AND the paragraph text content
    # is intact (the deprecation marking did not corrupt the text). (0.3 points)
    # Compound check: both conditions must hold simultaneously.
    # FAILS on initial_env (because strike is not True on the run) → PASSES on golden_env
    try:
        runs_with_text = [run for run in last_para.runs if run.text.strip()]
        last_para_text = last_para.text

        # Check 1: last paragraph has strike=True on at least one run (anchors to task change)
        has_any_strike = any(run.font.strike is True for run in runs_with_text)

        # Check 2: expected content fragments are present (text integrity)
        all_fragments_present = all(frag in last_para_text for frag in EXPECTED_LAST_PARA_FRAGMENTS)

        if has_any_strike and all_fragments_present:
            print(f"PASS: Component 2 — Last paragraph has strikethrough AND all expected text fragments are present (0.3 pts)")
            total_score += 0.3
        elif not has_any_strike:
            print(f"FAIL: Component 2 — Last paragraph does not have strikethrough formatting applied (prerequisite for text integrity check)")
        else:
            missing = [f for f in EXPECTED_LAST_PARA_FRAGMENTS if f not in last_para_text]
            print(f"FAIL: Component 2 — Missing expected text fragments in last paragraph: {missing}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
