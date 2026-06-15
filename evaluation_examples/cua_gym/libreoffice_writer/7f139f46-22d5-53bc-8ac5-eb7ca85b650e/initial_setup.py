"""
Initial Setup: Spanish verb flashcard table with 36 words in default black text
Task ID: osworld_writer_colorword_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_colorword_006'
OUTPUT = f'{WORKDIR}/spanish_verbs.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('Spanish Verb Flashcards', level=1)

    # Subtitle
    subtitle = doc.add_paragraph('Color-coded vocabulary reference for language learners.')

    # 36 Spanish words: mix of infinitives (ending -ar/-er/-ir) and conjugated forms
    # Arranged 6 rows x 6 columns
    words = [
        # Row 1
        'hablar', 'hablo', 'comer', 'como', 'vivir', 'vivo',
        # Row 2
        'correr', 'corro', 'bailar', 'bailo', 'escribir', 'escribo',
        # Row 3
        'cantar', 'canto', 'leer', 'leo', 'abrir', 'abro',
        # Row 4
        'caminar', 'camino', 'beber', 'bebo', 'salir', 'salgo',
        # Row 5
        'trabajar', 'trabajo', 'vender', 'vendo', 'subir', 'subo',
        # Row 6
        'llegar', 'llego', 'comprender', 'comprendo', 'partir', 'parto',
    ]

    # Create 6x6 table
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'

    for row_idx in range(6):
        for col_idx in range(6):
            word_idx = row_idx * 6 + col_idx
            cell = table.cell(row_idx, col_idx)
            # Clear default paragraph and set word with default black text
            para = cell.paragraphs[0]
            run = para.add_run(words[word_idx])
            run.font.size = Pt(12)
            # Explicitly set to black (default) — NO color pre-applied
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
