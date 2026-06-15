"""
Reward Script: Color Spanish verb infinitives navy blue and others orange in a word table
Task ID: osworld_writer_colorword_006
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All infinitives (words ending in -ar, -er, -ir) are colored navy blue (000080)
  Component 2 (0.3 pts): All non-infinitives are colored orange (FFA500)
  Component 3 (0.2 pts): No word remains in default black — every cell has been re-colored
"""

import os
from docx import Document
from docx.shared import RGBColor
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_colorword_006'

# Target colors
NAVY_BLUE = RGBColor(0x00, 0x00, 0x80)   # 000080
ORANGE    = RGBColor(0xFF, 0xA5, 0x00)   # FFA500

# Tolerance for color matching (Euclidean RGB distance)
COLOR_TOLERANCE = 50


def color_distance(c1, c2):
    """Compute Euclidean distance between two RGBColor values."""
    return sqrt(sum((int(a) - int(b)) ** 2 for a, b in zip(c1, c2)))


def is_infinitive(word):
    """Return True if word ends with -ar, -er, or -ir (infinitive marker)."""
    w = word.lower().strip()
    return w.endswith('ar') or w.endswith('er') or w.endswith('ir')


def get_cell_color(cell):
    """
    Return the first non-None RGBColor found in the cell runs, or None.
    """
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text.strip():
                try:
                    rgb = run.font.color.rgb
                    if rgb is not None:
                        return rgb
                except Exception:
                    pass
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must load
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: exactly 1 table with 6 rows x 6 cols (36 cells)
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    cells_data = []  # (word, color_or_None)
    for row in table.rows:
        for cell in row.cells:
            word = cell.text.strip()
            color = get_cell_color(cell)
            cells_data.append((word, color))

    if len(cells_data) == 0:
        print("CRITICAL: Table has no cells")
        print("REWARD: 0.0")
        return 0.0

    # Classify cells
    infinitives = [(w, c) for w, c in cells_data if is_infinitive(w) and w]
    non_infinitives = [(w, c) for w, c in cells_data if not is_infinitive(w) and w]

    print(f"Total cells: {len(cells_data)}")
    print(f"Infinitives found: {len(infinitives)} -> {[w for w, _ in infinitives]}")
    print(f"Non-infinitives found: {len(non_infinitives)} -> {[w for w, _ in non_infinitives]}")

    # ------------------------------------------------------------
    # Component 1: All infinitives are colored navy blue (0.5 pts)
    # ------------------------------------------------------------
    try:
        navy_pass_count = 0
        navy_fail_list = []
        for word, color in infinitives:
            if color is not None and color_distance(color, NAVY_BLUE) < COLOR_TOLERANCE:
                navy_pass_count += 1
            else:
                navy_fail_list.append((word, str(color) if color else "None"))

        if len(infinitives) > 0 and navy_pass_count == len(infinitives):
            print(f"PASS: Component 1 — All {len(infinitives)} infinitives are navy blue (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Only {navy_pass_count}/{len(infinitives)} infinitives are navy blue")
            if navy_fail_list:
                print(f"  Failed: {navy_fail_list[:5]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: All non-infinitives are colored orange (0.3 pts)
    # ---------------------------------------------------------------
    try:
        orange_pass_count = 0
        orange_fail_list = []
        for word, color in non_infinitives:
            if color is not None and color_distance(color, ORANGE) < COLOR_TOLERANCE:
                orange_pass_count += 1
            else:
                orange_fail_list.append((word, str(color) if color else "None"))

        if len(non_infinitives) > 0 and orange_pass_count == len(non_infinitives):
            print(f"PASS: Component 2 — All {len(non_infinitives)} non-infinitives are orange (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Only {orange_pass_count}/{len(non_infinitives)} non-infinitives are orange")
            if orange_fail_list:
                print(f"  Failed: {orange_fail_list[:5]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------
    # Component 3: No word remains in default black (000000) (0.2 pts)
    # This confirms the task was applied to ALL cells with no omissions.
    # -------------------------------------------------------------------
    try:
        BLACK = RGBColor(0x00, 0x00, 0x00)
        black_cells = []
        for word, color in cells_data:
            if not word:
                continue
            # A cell is "still black" if color is None (theme default black) or is 000000
            if color is None or color_distance(color, BLACK) < COLOR_TOLERANCE:
                black_cells.append((word, str(color) if color else "None/inherited"))

        if len(black_cells) == 0:
            print(f"PASS: Component 3 — No words remain in default black (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — {len(black_cells)} words still in black: {black_cells[:5]}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/spanish_verbs.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
