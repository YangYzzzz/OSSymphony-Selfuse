"""
Reward Script: Avery 5160 address labels (30 labels, 3x10 grid)
Task ID: writer_lec_041
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.4): Table exists with 10 rows x 3 cols
  - Component 2 (0.4): All 30 cells contain the correct full address
  - Component 3 (0.2): Each label has all 3 address lines
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_041'

EXPECTED_ADDRESS_LINES = [
    "TechStart Inc.",
    "200 Innovation Way",
    "San Jose, CA 95110",
]
EXPECTED_ROWS = 10
EXPECTED_COLS = 3


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with correct dimensions (0.4 points)
    # Initial file has 0 tables; golden has a 10x3 table
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 1 — No tables found in document")
        else:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_rows == EXPECTED_ROWS and num_cols == EXPECTED_COLS:
                print(f"PASS: Component 1 — Table is {num_rows}x{num_cols} (0.4 pts)")
                total_score += 0.4
            elif num_rows >= 8 and num_cols == EXPECTED_COLS:
                print(f"PARTIAL: Component 1 — Table is {num_rows}x{num_cols}, expected {EXPECTED_ROWS}x{EXPECTED_COLS} (0.2 pts)")
                total_score += 0.2  # partial credit
            else:
                print(f"FAIL: Component 1 — Table is {num_rows}x{num_cols}, expected {EXPECTED_ROWS}x{EXPECTED_COLS}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All 30 cells contain the correct full address (0.4 points)
    # Each cell should have "TechStart Inc.\n200 Innovation Way\nSan Jose, CA 95110"
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 2 — No table to check")
        else:
            table = doc.tables[0]
            expected_full = "\n".join(EXPECTED_ADDRESS_LINES)
            correct_count = 0
            total_cells = 0
            for row in table.rows:
                for cell in row.cells:
                    total_cells += 1
                    cell_text = cell.text.strip()
                    # Normalize whitespace for comparison
                    cell_normalized = "\n".join(line.strip() for line in cell_text.split("\n") if line.strip())
                    if cell_normalized == expected_full:
                        correct_count += 1

            if total_cells > 0 and correct_count == total_cells and total_cells == 30:
                print(f"PASS: Component 2 — All {correct_count}/{total_cells} cells have correct address (0.4 pts)")
                total_score += 0.4
            elif correct_count > 0:
                partial = round(0.4 * correct_count / max(total_cells, 30), 2)
                if partial > 0:
                    print(f"PARTIAL: Component 2 — {correct_count}/{total_cells} cells correct ({partial} pts)")
                    total_score += partial
            else:
                print(f"FAIL: Component 2 — {correct_count}/{total_cells} cells have correct address")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Each label has all 3 address lines (0.2 points)
    # Checks that each non-empty cell has exactly 3 lines matching the expected content
    try:
        if len(doc.tables) == 0:
            print(f"FAIL: Component 3 — No table to check")
        else:
            table = doc.tables[0]
            labels_with_3_lines = 0
            total_cells = 0
            for row in table.rows:
                for cell in row.cells:
                    total_cells += 1
                    lines = [line.strip() for line in cell.text.split("\n") if line.strip()]
                    if len(lines) == 3:
                        # Check each line contains key parts
                        has_company = "techstart" in lines[0].lower()
                        has_street = "innovation" in lines[1].lower() and "200" in lines[1]
                        has_city = "san jose" in lines[2].lower() and "95110" in lines[2]
                        if has_company and has_street and has_city:
                            labels_with_3_lines += 1

            if total_cells > 0 and labels_with_3_lines == total_cells and total_cells == 30:
                print(f"PASS: Component 3 — All {labels_with_3_lines} labels have 3 correct address lines (0.2 pts)")
                total_score += 0.2
            elif labels_with_3_lines > 0:
                partial = round(0.2 * labels_with_3_lines / max(total_cells, 30), 2)
                if partial > 0:
                    print(f"PARTIAL: Component 3 — {labels_with_3_lines}/{total_cells} labels have 3 lines ({partial} pts)")
                    total_score += partial
            else:
                print(f"FAIL: Component 3 — No labels have the expected 3 address lines")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
