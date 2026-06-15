"""
Reward Script: Configure body paragraphs with 1.5 line spacing, justified alignment,
               0.8 cm first-line indent, and 6pt space after
Task ID: writer_para_078
Domain: libreoffice_writer
Scoring:
  Component 1: All 5 body paragraphs have line_spacing=1.5          (0.25 pts)
  Component 2: All 5 body paragraphs have alignment=JUSTIFY         (0.25 pts)
  Component 3: All 5 body paragraphs have first_line_indent ~0.8cm  (0.25 pts)
  Component 4: All 5 body paragraphs have space_after=6pt           (0.25 pts)
  Total: 1.0
"""

import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_078'

# Body paragraph indices (0-based) in the document:
#   Para 0: "Annual Sustainability Report 2024" (Heading 1) — skip
#   Para 1: "Environmental Stewardship" (Heading 2) — skip
#   Para 2: body paragraph 1
#   Para 3: body paragraph 2
#   Para 4: "Social Responsibility" (Heading 2) — skip
#   Para 5: body paragraph 3
#   Para 6: body paragraph 4
#   Para 7: "Governance" (Heading 2) — skip
#   Para 8: body paragraph 5
BODY_PARA_INDICES = [2, 3, 5, 6, 8]

# Tolerance for first_line_indent comparison (±0.05 cm in EMU units)
INDENT_TOLERANCE_EMU = int(914400 * 2.54 * 0.05)  # ~116129 EMU (~0.05 cm)

# Expected values
EXPECTED_LINE_SPACING = 1.5
EXPECTED_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
EXPECTED_FIRST_INDENT_EMU = int(Cm(0.8))   # 288000 EMU
EXPECTED_SPACE_AFTER_PT = 6.0


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition gate: load the document
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify expected paragraph count
    if len(doc.paragraphs) < 9:
        print(f"CRITICAL: Expected at least 9 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Retrieve body paragraphs
    body_paras = [doc.paragraphs[i] for i in BODY_PARA_INDICES]

    # ----- Component 1: line_spacing == 1.5 for all body paragraphs (0.25 pts) -----
    try:
        ls_pass_count = 0
        ls_fail_details = []
        for idx, para in zip(BODY_PARA_INDICES, body_paras):
            pf = para.paragraph_format
            ls = pf.line_spacing
            if ls is not None and abs(float(ls) - EXPECTED_LINE_SPACING) < 0.01:
                ls_pass_count += 1
            else:
                ls_fail_details.append(f"Para[{idx}] line_spacing={ls!r}")

        if ls_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 1 — All {len(BODY_PARA_INDICES)} body paragraphs have line_spacing=1.5 (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — {ls_pass_count}/{len(BODY_PARA_INDICES)} body paragraphs have line_spacing=1.5")
            for detail in ls_fail_details:
                print(f"  FAIL detail: {detail}")
    except Exception as e:
        print(f"ERROR: Component 1 (line_spacing check) — {e}")

    # ----- Component 2: alignment == JUSTIFY for all body paragraphs (0.25 pts) -----
    try:
        align_pass_count = 0
        align_fail_details = []
        for idx, para in zip(BODY_PARA_INDICES, body_paras):
            pf = para.paragraph_format
            align = pf.alignment
            if align == EXPECTED_ALIGNMENT:
                align_pass_count += 1
            else:
                align_fail_details.append(f"Para[{idx}] alignment={align!r}")

        if align_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 2 — All {len(BODY_PARA_INDICES)} body paragraphs have alignment=JUSTIFY (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — {align_pass_count}/{len(BODY_PARA_INDICES)} body paragraphs have alignment=JUSTIFY")
            for detail in align_fail_details:
                print(f"  FAIL detail: {detail}")
    except Exception as e:
        print(f"ERROR: Component 2 (alignment check) — {e}")

    # ----- Component 3: first_line_indent ~0.8 cm for all body paragraphs (0.25 pts) -----
    try:
        indent_pass_count = 0
        indent_fail_details = []
        for idx, para in zip(BODY_PARA_INDICES, body_paras):
            pf = para.paragraph_format
            first_indent = pf.first_line_indent
            if (first_indent is not None and
                    abs(first_indent - EXPECTED_FIRST_INDENT_EMU) <= INDENT_TOLERANCE_EMU):
                indent_pass_count += 1
            else:
                # Convert to cm for readable output
                if first_indent is not None:
                    indent_cm = first_indent * 2.54 / 914400
                    indent_desc = f"{indent_cm:.4f} cm ({first_indent} EMU)"
                else:
                    indent_desc = "None"
                indent_fail_details.append(f"Para[{idx}] first_line_indent={indent_desc}")

        if indent_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 3 — All {len(BODY_PARA_INDICES)} body paragraphs have first_line_indent~0.8cm (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 3 — {indent_pass_count}/{len(BODY_PARA_INDICES)} body paragraphs have first_line_indent~0.8cm")
            for detail in indent_fail_details:
                print(f"  FAIL detail: {detail}")
    except Exception as e:
        print(f"ERROR: Component 3 (first_line_indent check) — {e}")

    # ----- Component 4: space_after == 6pt for all body paragraphs (0.25 pts) -----
    try:
        space_pass_count = 0
        space_fail_details = []
        for idx, para in zip(BODY_PARA_INDICES, body_paras):
            pf = para.paragraph_format
            space_after = pf.space_after
            if space_after is not None:
                space_after_pt = space_after / 12700  # EMU to pt
                if abs(space_after_pt - EXPECTED_SPACE_AFTER_PT) < 0.5:
                    space_pass_count += 1
                else:
                    space_fail_details.append(f"Para[{idx}] space_after={space_after_pt:.2f}pt ({space_after} EMU)")
            else:
                space_fail_details.append(f"Para[{idx}] space_after=None")

        if space_pass_count == len(BODY_PARA_INDICES):
            print(f"PASS: Component 4 — All {len(BODY_PARA_INDICES)} body paragraphs have space_after=6pt (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 — {space_pass_count}/{len(BODY_PARA_INDICES)} body paragraphs have space_after=6pt")
            for detail in space_fail_details:
                print(f"  FAIL detail: {detail}")
    except Exception as e:
        print(f"ERROR: Component 4 (space_after check) — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/sustainability_report.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
