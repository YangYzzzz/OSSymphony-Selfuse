"""
Initial Setup: Configure body paragraph formatting in a sustainability report
Task ID: writer_para_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_para_078'
OUTPUT = f'{WORKDIR}/sustainability_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Paragraph 1: Heading 1, center-aligned ---
    heading1 = doc.add_heading('Annual Sustainability Report 2024', level=1)
    heading1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Paragraph 2: Heading 2 ---
    doc.add_heading('Environmental Stewardship', level=2)

    # --- Paragraph 3: Body paragraph (no special formatting) ---
    p3 = doc.add_paragraph(
        'Our organization has made significant strides in reducing its environmental footprint '
        'over the past fiscal year. Total greenhouse gas emissions decreased by 18% compared to '
        'the 2023 baseline, exceeding our target of 15%.'
    )
    # No special formatting applied — left-aligned, default spacing

    # --- Paragraph 4: Body paragraph ---
    p4 = doc.add_paragraph(
        'Water consumption was reduced by 12% through the implementation of closed-loop cooling '
        'systems at our three largest manufacturing facilities. This initiative alone saved '
        'approximately 45 million liters of water.'
    )

    # --- Paragraph 5: Heading 2 ---
    doc.add_heading('Social Responsibility', level=2)

    # --- Paragraph 6: Body paragraph ---
    p6 = doc.add_paragraph(
        'We invested $8.5 million in community development programs across the 15 municipalities '
        'where we operate. Programs included STEM education grants, vocational training '
        'partnerships, and affordable housing initiatives.'
    )

    # --- Paragraph 7: Body paragraph ---
    p7 = doc.add_paragraph(
        'Employee volunteer hours reached a record 42,000 hours, with 78% of staff participating '
        'in at least one community service activity during the year.'
    )

    # --- Paragraph 8: Heading 2 ---
    doc.add_heading('Governance', level=2)

    # --- Paragraph 9: Body paragraph ---
    p9 = doc.add_paragraph(
        'The Board of Directors expanded the ESG Committee charter to include oversight of supply '
        'chain sustainability practices. A dedicated Chief Sustainability Officer was appointed in '
        'June 2024.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
