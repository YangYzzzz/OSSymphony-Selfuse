"""
Reward Script: Nested numbered list in table cell B2
Task ID: writer_lec_033
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Cell B2 has multiple paragraphs (split from single paragraph with line breaks)
  Component 2 (0.20): All paragraphs in B2 have numbering (numPr element present)
  Component 3 (0.25): Level-0 paragraphs use decimal (Arabic numeral) numbering format
  Component 4 (0.25): Level-1 paragraphs use lowerLetter numbering format
  Component 5 (0.10): Table structure preserved (3x3, other cells unchanged)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_033'
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_num_format(doc, num_id, ilvl_val):
    """Look up the numFmt for a given numId and ilvl in the numbering definitions."""
    try:
        numbering_part = doc.part.numbering_part
        numbering_xml = numbering_part._element
        ns = {'w': WNS}

        # Find the <w:num> with matching numId to get abstractNumId
        for num_elem in numbering_xml.findall('.//w:num', ns):
            nid = num_elem.get(f'{{{WNS}}}numId')
            if nid == str(num_id):
                abstract_ref = num_elem.find('w:abstractNumId', ns)
                if abstract_ref is None:
                    continue
                abstract_id = abstract_ref.get(f'{{{WNS}}}val')

                # Find the matching abstractNum
                for abn in numbering_xml.findall('.//w:abstractNum', ns):
                    abn_id = abn.get(f'{{{WNS}}}abstractNumId')
                    if abn_id == abstract_id:
                        for lvl in abn.findall('w:lvl', ns):
                            lvl_ilvl = lvl.get(f'{{{WNS}}}ilvl')
                            if lvl_ilvl == str(ilvl_val):
                                num_fmt = lvl.find('w:numFmt', ns)
                                if num_fmt is not None:
                                    return num_fmt.get(f'{{{WNS}}}val')
        return None
    except Exception:
        return None


def get_para_numbering(para):
    """Extract numId and ilvl from a paragraph's numPr, or (None, None)."""
    ns = {'w': WNS}
    pPr = para._element.find('w:pPr', ns)
    if pPr is None:
        return None, None
    numPr = pPr.find('w:numPr', ns)
    if numPr is None:
        return None, None
    ilvl_elem = numPr.find('w:ilvl', ns)
    numId_elem = numPr.find('w:numId', ns)
    ilvl = ilvl_elem.get(f'{{{WNS}}}val') if ilvl_elem is not None else None
    numId = numId_elem.get(f'{{{WNS}}}val') if numId_elem is not None else None
    return numId, ilvl


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: Must have at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: Table must be at least 2x2
    if len(table.rows) < 2 or len(table.columns) < 2:
        print("CRITICAL: Table too small (need at least 2x2)")
        print("REWARD: 0.0")
        return 0.0

    cell_b2 = table.cell(1, 1)
    paras = cell_b2.paragraphs

    # Component 1: Cell B2 has multiple paragraphs (0.20 points)
    # Initial state has 1 paragraph with \n line breaks; golden has 7 separate paragraphs
    try:
        para_count = len(paras)
        if para_count >= 3:
            print(f"PASS: Component 1 - Cell B2 has {para_count} paragraphs (>= 3) (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 - Cell B2 has {para_count} paragraphs, expected >= 3")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: All paragraphs in B2 have numbering (0.20 points)
    # Initial state has no numPr on any paragraph; golden has numPr on all
    try:
        numbered_count = 0
        total_paras = 0
        for para in paras:
            if not para.text.strip():
                continue  # skip empty paragraphs
            total_paras += 1
            numId, ilvl = get_para_numbering(para)
            if numId is not None and numId != '0':
                numbered_count += 1
        if total_paras > 0 and numbered_count == total_paras:
            print(f"PASS: Component 2 - All {numbered_count}/{total_paras} paragraphs have numbering (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 - Only {numbered_count}/{total_paras} paragraphs have numbering")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Level-0 paragraphs use decimal (Arabic numeral) format (0.25 points)
    # Task requires main items use Arabic numbers (1., 2., 3.)
    try:
        level0_paras = []
        for para in paras:
            if not para.text.strip():
                continue
            numId, ilvl = get_para_numbering(para)
            if numId is not None and ilvl == '0':
                level0_paras.append((para.text, numId))

        if len(level0_paras) == 0:
            print("FAIL: Component 3 - No level-0 numbered paragraphs found")
        else:
            non_decimal = [t for t, nid in level0_paras if get_num_format(doc, nid, 0) != 'decimal']
            if len(non_decimal) == 0:
                print(f"PASS: Component 3 - All {len(level0_paras)} level-0 paragraphs use decimal format (0.25 pts)")
                total_score += 0.25
            else:
                for t in non_decimal:
                    print(f"  Level-0 para '{t[:40]}...' does not use decimal format")
                print(f"FAIL: Component 3 - {len(non_decimal)} level-0 paragraphs do not use decimal format")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Level-1 paragraphs use lowerLetter format (0.25 points)
    # Task requires sub-items use lowercase letters (a., b.)
    try:
        level1_paras = []
        for para in paras:
            if not para.text.strip():
                continue
            numId, ilvl = get_para_numbering(para)
            if numId is not None and ilvl == '1':
                level1_paras.append((para.text, numId))

        if len(level1_paras) == 0:
            print("FAIL: Component 4 - No level-1 numbered paragraphs found")
        else:
            non_lower = [t for t, nid in level1_paras if get_num_format(doc, nid, 1) != 'lowerLetter']
            if len(non_lower) == 0:
                print(f"PASS: Component 4 - All {len(level1_paras)} level-1 paragraphs use lowerLetter format (0.25 pts)")
                total_score += 0.25
            else:
                for t in non_lower:
                    print(f"  Level-1 para '{t[:40]}...' does not use lowerLetter format")
                print(f"FAIL: Component 4 - {len(non_lower)} level-1 paragraphs do not use lowerLetter format")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Table structure preserved AND B2 has numbering (0.10 points)
    # Compound check: table must remain 3x3 with other cells unchanged,
    # AND B2 must have numbering (anchors this to the task change)
    try:
        rows = len(table.rows)
        cols = len(table.columns)
        expected_cells = {
            (0, 0): 'Phase',
            (0, 1): 'Task Breakdown',
            (0, 2): 'Status',
            (1, 0): 'Phase 1: Discovery',
            (1, 2): 'In Progress',
            (2, 0): 'Phase 2: Design',
            (2, 2): 'Not Started',
        }
        structure_ok = (rows == 3 and cols == 3)
        if not structure_ok:
            print(f"FAIL: Component 5 - Table is {rows}x{cols}, expected 3x3")
        else:
            mismatched = [(r, c, expected_text, table.cell(r, c).text.strip())
                          for (r, c), expected_text in expected_cells.items()
                          if table.cell(r, c).text.strip() != expected_text]

            # Anchor to task change: B2 must also have numbering
            b2_numbered = any(
                get_para_numbering(p)[0] not in (None, '0')
                for p in paras if p.text.strip()
            )

            if len(mismatched) == 0 and b2_numbered:
                print(f"PASS: Component 5 - Table structure preserved with numbered B2 (0.10 pts)")
                total_score += 0.10
            elif len(mismatched) > 0:
                for r, c, exp, act in mismatched:
                    print(f"  Cell({r},{c}): expected '{exp}', found '{act}'")
                print(f"FAIL: Component 5 - Some cells have unexpected values")
            else:
                print(f"FAIL: Component 5 - Table structure OK but B2 has no numbering")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
