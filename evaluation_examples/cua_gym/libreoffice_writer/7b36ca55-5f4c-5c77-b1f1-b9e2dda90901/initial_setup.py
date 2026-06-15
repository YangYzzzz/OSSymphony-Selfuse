"""
Initial Setup: Create a Writer document with a 3x3 project tracking table.
Cell B2 contains plain text with task descriptions and sub-tasks separated by line breaks.
Task ID: writer_lec_033
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_033'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Website Redesign Project Plan", level=1)

    # Introductory paragraph
    doc.add_paragraph(
        "This document outlines the project tracking matrix for the upcoming "
        "website redesign initiative. All team leads should update their section "
        "weekly by Friday 5 PM EST."
    )

    # Create 3x3 project tracking table
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Row 0: Headers
    headers = ["Phase", "Task Breakdown", "Status"]
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Row 1 (data row index 1)
    # Cell A2 (0-indexed: row=1, col=0)
    table.cell(1, 0).text = "Phase 1: Discovery"

    # Cell B2 (0-indexed: row=1, col=1) — plain text with line breaks, NO list formatting
    cell_b2 = table.cell(1, 1)
    cell_b2.text = ""
    para = cell_b2.paragraphs[0]
    para.add_run("Conduct stakeholder interviews\n")
    para.add_run("Identify key pain points\n")
    para.add_run("Document current user journeys\n")
    para.add_run("Analyze competitor websites\n")
    para.add_run("Review existing analytics data\n")
    para.add_run("Compile research findings report\n")
    para.add_run("Present findings to leadership")

    # Cell C2 (row=1, col=2)
    table.cell(1, 2).text = "In Progress"

    # Row 2 (data row index 2)
    table.cell(2, 0).text = "Phase 2: Design"

    cell_b3 = table.cell(2, 1)
    cell_b3.text = "Create wireframes for homepage and landing pages. " \
                   "Develop interactive prototypes. Conduct usability testing sessions."

    table.cell(2, 2).text = "Not Started"

    # Additional context paragraph after table
    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Phase 1 task breakdown in column B should be reformatted as a "
        "structured numbered list with sub-items for better readability."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
