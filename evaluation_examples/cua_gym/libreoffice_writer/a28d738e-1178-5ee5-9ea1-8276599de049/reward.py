"""
FINAL REWARD SCRIPT - SUCCESS
Task: I pasted a CSV and LibreOffice automatically turned it into something called “Table 1.” Before I can upload the content to our database, I need that table flattened back into plain text with TAB characters between each column. What’s the quickest way to convert Table 1 to tab-separated text in Writer?
Generated: 2025-09-10 14:06:14
Status: success
Model: azure-o3
Total Steps: 2
"""

import os
from docx import Document


def verify_writer_flatten_table(file_path: str) -> float:
    """Verify that a LibreOffice Writer/Word document which originally
    contained a CSV-generated table has been flattened back into plain
    tab-separated text.

    Scoring (progressive – totals to 1.0):
    1. Table removal (0.4) – the document must contain *no* tables.
    2. Presence of TAB characters (0.3) – text must include at least one "\t".
    3. Column consistency (0.3) – every non-empty line must have the same
       number of TAB-separated columns (and at least 2 columns).
    """

    print(f"Verifying file: {file_path}\n")

    score = 0.0
    max_score = 1.0

    # ------------------------------------------------------------------
    # 0. Prerequisite – file exists & can be loaded (no points awarded)
    # ------------------------------------------------------------------
    if not os.path.exists(file_path):
        print("✗ File not found")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Failed to load DOCX: {exc}")
        return 0.0

    # ------------------------------------------------------------------
    # 1. Table removal (0.4 pts)
    # ------------------------------------------------------------------
    if len(doc.tables) == 0:
        print("✓ No tables present (requirement 1: table removed) – 0.4 pts")
        score += 0.4
    else:
        print(f"✗ Still found {len(doc.tables)} table(s) – 0 pts")

    # Gather all non-empty paragraph texts
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    # ------------------------------------------------------------------
    # 2. Presence of TAB characters (0.3 pts)
    # ------------------------------------------------------------------
    if "\t" in full_text:
        print("✓ Tab characters detected in text (requirement 2) – 0.3 pts")
        score += 0.3
    else:
        print("✗ No tab characters found – 0 pts")

    # ------------------------------------------------------------------
    # 3. Column consistency across lines (0.3 pts)
    # ------------------------------------------------------------------
    lines = [ln for ln in full_text.split("\n") if ln.strip()]
    if lines:
        column_counts = [len(ln.split("\t")) for ln in lines]
        expected_cols = column_counts[0]
        if expected_cols >= 2 and all(c == expected_cols for c in column_counts):
            print(
                f"✓ All {len(lines)} line(s) have consistent column count "
                f"({expected_cols}) (requirement 3) – 0.3 pts"
            )
            score += 0.3
        else:
            print("✗ Inconsistent column counts or fewer than 2 columns – 0 pts")
    else:
        print("✗ Document contains no non-empty lines to evaluate – 0 pts")

    # ------------------------------------------------------------------
    final_score = min(score, max_score)
    print(f"\nFinal score: {final_score}")
    return final_score


if __name__ == "__main__":
    FILE_PATH = (
        "/home/user/i_pasted_a_csv_and_libreoffice_automatically_" \
        "turned_it_into_something_called_table_1_before_i_can_up.docx"
    )
    reward = verify_writer_flatten_table(FILE_PATH)
    print(f"REWARD: {reward}")
