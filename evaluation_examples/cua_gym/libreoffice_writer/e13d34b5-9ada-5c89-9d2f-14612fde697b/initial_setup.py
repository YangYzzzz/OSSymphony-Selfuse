"""
Initial Setup: Reorder paragraphs in a recipe document
Task ID: wrpara_023
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_023'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Title
    title = doc.add_heading("Grandma's Classic Beef Stew", level=1)

    # Paragraph 2: Ingredients list
    p2 = doc.add_paragraph()
    run_h = p2.add_run("Ingredients\n")
    run_h.bold = True
    run_h.font.size = Pt(13)
    p2.add_run(
        "2 lbs beef chuck, cut into 1-inch cubes\n"
        "3 tablespoons olive oil\n"
        "1 large yellow onion, diced\n"
        "4 cloves garlic, minced\n"
        "4 medium Yukon Gold potatoes, quartered\n"
        "3 large carrots, sliced into rounds\n"
        "2 stalks celery, chopped\n"
        "1 cup frozen peas\n"
        "4 cups beef broth\n"
        "1 cup dry red wine\n"
        "2 tablespoons tomato paste\n"
        "1 teaspoon dried thyme\n"
        "2 bay leaves\n"
        "Salt and freshly ground black pepper to taste\n"
        "3 tablespoons all-purpose flour"
    )

    # Paragraph 3: Equipment needed
    p3 = doc.add_paragraph()
    run_h3 = p3.add_run("Equipment Needed\n")
    run_h3.bold = True
    run_h3.font.size = Pt(13)
    p3.add_run(
        "Large Dutch oven or heavy-bottomed pot (5-quart minimum)\n"
        "Sharp chef's knife and cutting board\n"
        "Wooden spoon or heat-resistant spatula\n"
        "Measuring cups and spoons\n"
        "Ladle for serving\n"
        "Paper towels for patting beef dry"
    )

    # Paragraph 4: Overview/Summary
    p4 = doc.add_paragraph()
    run_h4 = p4.add_run("Overview\n")
    run_h4.bold = True
    run_h4.font.size = Pt(13)
    p4.add_run(
        "This hearty beef stew is a family recipe passed down through three generations. "
        "Tender chunks of beef are slow-simmered with root vegetables in a rich, wine-infused broth "
        "until everything is melt-in-your-mouth tender. The secret is patience: low and slow cooking "
        "transforms tough cuts of meat into a comforting, deeply flavored dish that's perfect for cold "
        "winter evenings. Total preparation and cooking time is approximately 2.5 hours, with most of "
        "that being hands-off simmering."
    )

    # Paragraph 5: Instructions
    p5 = doc.add_paragraph()
    run_h5 = p5.add_run("Instructions\n")
    run_h5.bold = True
    run_h5.font.size = Pt(13)
    p5.add_run(
        "1. Pat the beef cubes dry with paper towels and season generously with salt, pepper, and flour.\n"
        "2. Heat olive oil in the Dutch oven over medium-high heat until shimmering.\n"
        "3. Brown the beef in batches, about 3 minutes per side. Do not overcrowd the pot. Transfer to a plate.\n"
        "4. Reduce heat to medium. Add onion and cook until softened, about 5 minutes.\n"
        "5. Add garlic and tomato paste, stirring constantly for 1 minute until fragrant.\n"
        "6. Pour in the red wine, scraping the bottom of the pot to deglaze.\n"
        "7. Return beef to the pot. Add broth, thyme, and bay leaves. Bring to a boil.\n"
        "8. Reduce heat to low, cover, and simmer for 1 hour.\n"
        "9. Add potatoes, carrots, and celery. Continue simmering covered for 45 minutes.\n"
        "10. Add frozen peas during the last 10 minutes of cooking.\n"
        "11. Remove bay leaves. Adjust seasoning with salt and pepper to taste."
    )

    # Paragraph 6: Serving suggestions
    p6 = doc.add_paragraph()
    run_h6 = p6.add_run("Serving Suggestions\n")
    run_h6.bold = True
    run_h6.font.size = Pt(13)
    p6.add_run(
        "Ladle the stew into deep bowls and serve with crusty sourdough bread for dipping. "
        "A dollop of sour cream on top adds a lovely tangy contrast. For a complete meal, pair with "
        "a simple green salad dressed in lemon vinaigrette. This stew tastes even better the next day "
        "as the flavors continue to meld. Leftovers can be refrigerated for up to 4 days or frozen "
        "for up to 3 months. Reheat gently on the stovetop, adding a splash of broth if needed."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
