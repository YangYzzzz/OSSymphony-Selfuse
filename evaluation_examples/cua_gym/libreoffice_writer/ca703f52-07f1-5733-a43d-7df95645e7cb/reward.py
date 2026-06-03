"""
Reward Script: Find and replace 'Fiscal Year 2023' with 'Fiscal Year 2024' everywhere
Task ID: writer_edit_045
Domain: libreoffice_writer
Scoring:
  Component 1 (0.40): Body paragraphs — all body text occurrences replaced (3 of 3)
  Component 2 (0.20): Header — header text updated to 'Fiscal Year 2024'
  Component 3 (0.20): Footer — footer text updated to 'Fiscal Year 2024'
  Component 4 (0.20): Text box — text box on page 2 updated to 'Fiscal Year 2024'
  Total: 1.0

Ground truth from context:
  - 'Fiscal Year 2023' appears 6 times in initial file:
      3 times in body text, 1 in page header, 1 in text box (page 2), 1 in page footer
  - After task: all 6 occurrences must say 'Fiscal Year 2024'
  - File: /home/user/Desktop/financial_summary.docx
"""

import os
import re

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_045'
FILE_PATH = '/home/user/Desktop/financial_summary.docx'

OLD_TEXT = 'Fiscal Year 2023'
NEW_TEXT = 'Fiscal Year 2024'

# Expected body paragraph count with FY2024 after task
EXPECTED_BODY_COUNT = 3


def count_in_paragraphs(paragraphs, search_text):
    """Count occurrences of search_text across a list of paragraphs (using para.text)."""
    count = 0
    for para in paragraphs:
        # para.text concatenates all runs, safe for search purposes
        count += para.text.count(search_text)
    return count


def count_in_textboxes(doc, search_text):
    """Count occurrences of search_text inside text boxes (txbxContent) via XML."""
    body = doc.element.body
    # Serialize to string for regex search — avoids lxml dependency
    import zipfile
    # Use lxml via element tree
    from lxml import etree
    xml_str = etree.tostring(body, encoding='unicode')
    # Extract text from txbxContent blocks
    textbox_blocks = re.findall(
        r'<w:txbxContent\b[^>]*>.*?</w:txbxContent>',
        xml_str, re.DOTALL
    )
    count = 0
    for block in textbox_blocks:
        # Extract all <w:t> node contents
        texts = re.findall(r'<w:t(?:\s[^>]*)?>(.*?)</w:t>', block)
        combined = ' '.join(texts)
        count += combined.count(search_text)
    return count


def verify_task(file_path):
    """
    Verify that all 6 occurrences of 'Fiscal Year 2023' have been replaced
    with 'Fiscal Year 2024' in the document, including headers, footers, and text boxes.

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # --- Load document ---
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # === Component 1: Body text replacements (0.40 points) ===
    # The task requires changing 3 body-paragraph occurrences of 'Fiscal Year 2023'
    # to 'Fiscal Year 2024'. After completion, OLD_TEXT must not appear in body
    # paragraphs, and NEW_TEXT must appear exactly 3 times there.
    try:
        old_in_body = count_in_paragraphs(doc.paragraphs, OLD_TEXT)
        new_in_body = count_in_paragraphs(doc.paragraphs, NEW_TEXT)

        if old_in_body == 0 and new_in_body >= EXPECTED_BODY_COUNT:
            print(f"PASS: Component 1 — body paragraphs: "
                  f"{new_in_body} occurrence(s) of '{NEW_TEXT}' found, "
                  f"0 residual '{OLD_TEXT}' (0.40 pts)")
            total_score += 0.40
        else:
            print(f"FAIL: Component 1 — body paragraphs: "
                  f"found {old_in_body} residual '{OLD_TEXT}', "
                  f"{new_in_body} '{NEW_TEXT}' (need {EXPECTED_BODY_COUNT})")
    except Exception as e:
        print(f"ERROR: Component 1 — body paragraphs: {e}")

    # === Component 2: Header replacement (0.20 points) ===
    # The task requires updating the page header which contained 'Fiscal Year 2023'.
    try:
        header_old_count = 0
        header_new_count = 0
        for section in doc.sections:
            header_old_count += count_in_paragraphs(section.header.paragraphs, OLD_TEXT)
            header_new_count += count_in_paragraphs(section.header.paragraphs, NEW_TEXT)

        if header_old_count == 0 and header_new_count >= 1:
            print(f"PASS: Component 2 — page header: "
                  f"'{NEW_TEXT}' found in header, 0 residual '{OLD_TEXT}' (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — page header: "
                  f"found {header_old_count} residual '{OLD_TEXT}', "
                  f"{header_new_count} '{NEW_TEXT}'")
    except Exception as e:
        print(f"ERROR: Component 2 — page header: {e}")

    # === Component 3: Footer replacement (0.20 points) ===
    # The task requires updating the page footer which contained 'Fiscal Year 2023'.
    try:
        footer_old_count = 0
        footer_new_count = 0
        for section in doc.sections:
            footer_old_count += count_in_paragraphs(section.footer.paragraphs, OLD_TEXT)
            footer_new_count += count_in_paragraphs(section.footer.paragraphs, NEW_TEXT)

        if footer_old_count == 0 and footer_new_count >= 1:
            print(f"PASS: Component 3 — page footer: "
                  f"'{NEW_TEXT}' found in footer, 0 residual '{OLD_TEXT}' (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 3 — page footer: "
                  f"found {footer_old_count} residual '{OLD_TEXT}', "
                  f"{footer_new_count} '{NEW_TEXT}'")
    except Exception as e:
        print(f"ERROR: Component 3 — page footer: {e}")

    # === Component 4: Text box replacement (0.20 points) ===
    # The task requires updating the text box on page 2 which contained 'Fiscal Year 2023'.
    # Text boxes are stored as txbxContent in the document XML, not accessible via
    # doc.paragraphs — we must parse the raw XML.
    try:
        from lxml import etree
        body = doc.element.body
        xml_str = etree.tostring(body, encoding='unicode')

        textbox_blocks = re.findall(
            r'<w:txbxContent\b[^>]*>.*?</w:txbxContent>',
            xml_str, re.DOTALL
        )

        tb_old_count = 0
        tb_new_count = 0
        for block in textbox_blocks:
            texts = re.findall(r'<w:t(?:\s[^>]*)?>(.*?)</w:t>', block)
            combined = ' '.join(texts)
            tb_old_count += combined.count(OLD_TEXT)
            tb_new_count += combined.count(NEW_TEXT)

        if tb_old_count == 0 and tb_new_count >= 1:
            print(f"PASS: Component 4 — text box: "
                  f"'{NEW_TEXT}' found in text box(es), "
                  f"0 residual '{OLD_TEXT}' (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — text box: "
                  f"found {tb_old_count} residual '{OLD_TEXT}', "
                  f"{tb_new_count} '{NEW_TEXT}' in text boxes")
    except Exception as e:
        print(f"ERROR: Component 4 — text box: {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point — script runs on the VM where WORKDIR == /home/user
if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
