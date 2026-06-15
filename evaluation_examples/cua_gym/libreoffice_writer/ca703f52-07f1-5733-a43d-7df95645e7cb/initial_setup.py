"""
Initial Setup: Financial Summary Document with 'Fiscal Year 2023' occurrences
Task ID: writer_edit_045
Domain: libreoffice_writer

Creates a 4-page financial document with 'Fiscal Year 2023' appearing 6 times:
  - 3 times in body text
  - 1 time in page header
  - 1 time in text box on page 2
  - 1 time in page footer
"""

import os
import shlex
import subprocess
import time
from copy import deepcopy
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_045'
OUTPUT = f'{WORKDIR}/Desktop/financial_summary.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_text_box(paragraph, text, width_emu, height_emu, left_emu, top_emu):
    """
    Insert a drawing text box (inline shape) into the paragraph via raw OOXML.
    Uses wps:txbx (Word Processing Shapes) to create a floating text box.
    """
    # Namespace map
    nsmap = {
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    }

    # Build raw XML for drawing with text box
    drawing_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <wp:anchor distT="114300" distB="114300" distL="114300" distR="114300"
    simplePos="0" relativeHeight="251658240" behindDoc="0" locked="0"
    layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="column">
      <wp:posOffset>{left_emu}</wp:posOffset>
    </wp:positionH>
    <wp:positionV relativeFrom="paragraph">
      <wp:posOffset>{top_emu}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="1" name="TextBox 1"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr txBx="1">
            <a:spLocks noChangeArrowheads="1"/>
          </wps:cNvSpPr>
          <wps:spPr>
            <a:xfrm>
              <a:off x="{left_emu}" y="{top_emu}"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
          </wps:spPr>
          <wps:txbx>
            <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <w:p>
                <w:r>
                  <w:t xml:space="preserve">Fiscal Year 2023 — Quarterly Performance Summary</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </wps:txbx>
          <wps:bodyPr insFit="auto">
            <a:noAutofit/>
          </wps:bodyPr>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''

    drawing_element = etree.fromstring(drawing_xml)
    run = paragraph.add_run()
    run._element.append(drawing_element)


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Header: "Fiscal Year 2023" occurrence #1 ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    run_h = hp.add_run("Acme Corporation — Annual Report | Fiscal Year 2023")
    run_h.font.size = Pt(10)
    run_h.font.name = "Calibri"
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Footer: "Fiscal Year 2023" occurrence #2 ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    run_f = fp.add_run("Fiscal Year 2023 | Confidential — Acme Corporation | Page ")
    run_f.font.size = Pt(9)
    run_f.font.name = "Calibri"
    # Page number field
    r1 = fp.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1._element.append(fld_begin)
    r2 = fp.add_run()
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r3._element.append(fld_end)
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ===================== PAGE 1 =====================
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("ACME CORPORATION")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.name = "Calibri"

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Body occurrence #3: "Fiscal Year 2023" in title/subtitle
    sub_run = subtitle_para.add_run("Annual Financial Report — Fiscal Year 2023")
    sub_run.bold = True
    sub_run.font.size = Pt(16)
    sub_run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # Executive Summary heading
    h1 = doc.add_heading("Executive Summary", level=1)

    # Executive summary body text
    exec_para1 = doc.add_paragraph(
        "This report presents a comprehensive overview of Acme Corporation's financial performance "
        "for the fiscal year ending December 31, 2023. The company demonstrated strong resilience "
        "in a challenging macroeconomic environment, achieving record revenues across three of its "
        "four business units."
    )
    exec_para1.style.font.size = Pt(11)

    exec_para2 = doc.add_paragraph(
        "Total revenues for Fiscal Year 2023 reached $2.47 billion, representing a 12.3% increase "
        "compared to the prior fiscal year. Operating income improved by 18.7% to $485 million, "
        "driven primarily by operational efficiency initiatives and disciplined cost management."
    )
    # Body occurrence #4: "Fiscal Year 2023" in body text
    # (already embedded above as text in exec_para2)

    exec_para3 = doc.add_paragraph(
        "Key performance highlights include: (1) Revenue growth of 12.3% year-over-year; "
        "(2) Operating margin expansion of 110 basis points to 19.6%; "
        "(3) Free cash flow generation of $312 million; "
        "(4) Successful integration of two strategic acquisitions."
    )

    # Financial Highlights Table
    doc.add_heading("Financial Highlights", level=2)

    table1 = doc.add_table(rows=6, cols=3)
    table1.style = "Table Grid"

    headers_row = ["Metric", "FY 2023", "FY 2022"]
    for j, h_text in enumerate(headers_row):
        cell = table1.cell(0, j)
        run_cell = cell.paragraphs[0].add_run(h_text)
        run_cell.bold = True
        run_cell.font.size = Pt(11)

    data_rows = [
        ["Total Revenue", "$2.47B", "$2.20B"],
        ["Operating Income", "$485M", "$408M"],
        ["Net Income", "$312M", "$267M"],
        ["Earnings Per Share", "$4.82", "$4.11"],
        ["Free Cash Flow", "$312M", "$285M"],
    ]
    for i, row_data in enumerate(data_rows, 1):
        for j, val in enumerate(row_data):
            table1.cell(i, j).paragraphs[0].add_run(val).font.size = Pt(11)

    doc.add_page_break()

    # ===================== PAGE 2 =====================
    doc.add_heading("Segment Performance", level=1)

    seg_para1 = doc.add_paragraph(
        "The following section provides detailed analysis of each of Acme Corporation's four "
        "primary business segments and their contributions to the overall consolidated results."
    )

    # Text box on page 2: "Fiscal Year 2023" occurrence #5
    tb_para = doc.add_paragraph()
    add_text_box(
        tb_para,
        text="Fiscal Year 2023 — Quarterly Performance Summary",
        width_emu=int(Inches(4.5)),
        height_emu=int(Inches(0.8)),
        left_emu=int(Inches(1.0)),
        top_emu=int(Inches(0.2)),
    )

    doc.add_heading("Technology Solutions Division", level=2)
    tech_para = doc.add_paragraph(
        "The Technology Solutions Division recorded revenues of $892 million in the current fiscal "
        "year, an increase of 21.4% year-over-year. Growth was driven by strong demand for cloud "
        "migration services and cybersecurity consulting engagements. The division secured 47 new "
        "enterprise contracts with an average contract value of $12.3 million."
    )

    table2 = doc.add_table(rows=5, cols=3)
    table2.style = "Table Grid"
    t2_headers = ["Quarter", "Revenue ($M)", "YoY Growth"]
    for j, h_text in enumerate(t2_headers):
        run_cell = table2.cell(0, j).paragraphs[0].add_run(h_text)
        run_cell.bold = True
        run_cell.font.size = Pt(11)

    t2_data = [
        ["Q1 2023", "$198.4", "+18.2%"],
        ["Q2 2023", "$214.7", "+20.1%"],
        ["Q3 2023", "$231.8", "+23.4%"],
        ["Q4 2023", "$247.1", "+24.0%"],
    ]
    for i, row_data in enumerate(t2_data, 1):
        for j, val in enumerate(row_data):
            table2.cell(i, j).paragraphs[0].add_run(val).font.size = Pt(11)

    doc.add_heading("Industrial Manufacturing Division", level=2)
    indus_para = doc.add_paragraph(
        "Industrial Manufacturing revenue reached $634 million, reflecting modest growth of 4.7% "
        "versus the prior year period. The division faced headwinds from elevated raw material "
        "costs and supply chain disruptions, partially offset by price realization and productivity "
        "improvements at the Monterrey and Guangzhou facilities."
    )

    doc.add_page_break()

    # ===================== PAGE 3 =====================
    doc.add_heading("Financial Statements", level=1)

    # Body occurrence #5: "Fiscal Year 2023" in body text paragraph
    fs_para = doc.add_paragraph()
    fs_run = fs_para.add_run(
        "The consolidated financial statements for Fiscal Year 2023 have been prepared in "
        "accordance with U.S. Generally Accepted Accounting Principles (GAAP) and have been "
        "audited by Deloitte & Touche LLP. The following statements present the company's "
        "financial position as of December 31, 2023."
    )
    fs_run.font.size = Pt(11)

    doc.add_heading("Consolidated Income Statement", level=2)

    table3 = doc.add_table(rows=9, cols=3)
    table3.style = "Table Grid"
    t3_headers = ["Line Item", "FY 2023 ($M)", "FY 2022 ($M)"]
    for j, h_text in enumerate(t3_headers):
        run_cell = table3.cell(0, j).paragraphs[0].add_run(h_text)
        run_cell.bold = True
        run_cell.font.size = Pt(11)

    t3_data = [
        ["Total Revenue", "2,470", "2,199"],
        ["Cost of Goods Sold", "(1,482)", "(1,338)"],
        ["Gross Profit", "988", "861"],
        ["Operating Expenses", "(503)", "(453)"],
        ["Operating Income", "485", "408"],
        ["Interest Expense", "(42)", "(38)"],
        ["Income Before Tax", "443", "370"],
        ["Net Income", "312", "267"],
    ]
    for i, row_data in enumerate(t3_data, 1):
        for j, val in enumerate(row_data):
            table3.cell(i, j).paragraphs[0].add_run(val).font.size = Pt(11)

    doc.add_heading("Consolidated Balance Sheet", level=2)

    bs_para = doc.add_paragraph(
        "The balance sheet reflects a strong financial position with total assets of $4.82 billion "
        "and total equity of $2.63 billion as of December 31, 2023. Net debt decreased by "
        "$187 million to $891 million, reflecting strong free cash flow generation."
    )

    table4 = doc.add_table(rows=7, cols=3)
    table4.style = "Table Grid"
    t4_headers = ["Asset Category", "Dec 2023 ($M)", "Dec 2022 ($M)"]
    for j, h_text in enumerate(t4_headers):
        run_cell = table4.cell(0, j).paragraphs[0].add_run(h_text)
        run_cell.bold = True
        run_cell.font.size = Pt(11)

    t4_data = [
        ["Cash & Equivalents", "623", "511"],
        ["Accounts Receivable", "412", "387"],
        ["Inventories", "298", "314"],
        ["Property, Plant & Equipment", "1,847", "1,792"],
        ["Goodwill & Intangibles", "1,638", "1,612"],
        ["Total Assets", "4,818", "4,616"],
    ]
    for i, row_data in enumerate(t4_data, 1):
        for j, val in enumerate(row_data):
            table4.cell(i, j).paragraphs[0].add_run(val).font.size = Pt(11)

    doc.add_page_break()

    # ===================== PAGE 4 =====================
    doc.add_heading("Outlook and Forward Guidance", level=1)

    # Body occurrence #6: "Fiscal Year 2023" in body text paragraph
    outlook_para1 = doc.add_paragraph()
    outlook_run1 = outlook_para1.add_run(
        "Building on the strong performance achieved in Fiscal Year 2023, Acme Corporation's "
        "management team is confident in the company's trajectory for the upcoming fiscal year. "
        "The strategic initiatives launched during the year have established a solid foundation "
        "for continued growth and value creation for shareholders."
    )
    outlook_run1.font.size = Pt(11)

    outlook_para2 = doc.add_paragraph(
        "For the fiscal year ending December 31, 2024, management provides the following guidance: "
        "total revenue in the range of $2.70 to $2.80 billion, representing growth of 9-13% "
        "versus the prior year; operating income in the range of $540 to $560 million; and "
        "diluted earnings per share of $5.20 to $5.45."
    )

    doc.add_heading("Strategic Priorities", level=2)

    priorities = [
        "Expand Technology Solutions Division through organic growth and targeted acquisitions",
        "Optimize Industrial Manufacturing operations to improve margin by 150-200 basis points",
        "Launch next-generation product platform in the Consumer Products Division",
        "Pursue disciplined capital allocation with continued share repurchase program",
        "Advance ESG commitments with 30% reduction in Scope 1 and 2 emissions by 2030",
    ]
    for priority in priorities:
        doc.add_paragraph(priority, style="List Bullet")

    doc.add_heading("Risk Factors", level=2)
    risk_para = doc.add_paragraph(
        "As with all forward-looking statements, there are risks and uncertainties that could "
        "cause actual results to differ materially from management's guidance. Key risk factors "
        "include macroeconomic conditions, foreign exchange fluctuations, supply chain disruptions, "
        "regulatory changes, and competitive dynamics in our end markets. Investors are encouraged "
        "to review the full risk factor disclosure in the company's Form 10-K filing."
    )

    # Closing signature block
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    sig_run = sig_para.add_run(
        "Prepared by: Acme Corporation Finance Department\n"
        "Date: March 31, 2024"
    )
    sig_run.font.size = Pt(10)
    sig_run.italic = True

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
