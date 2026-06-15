"""
Initial Setup: Create a Writer document with an 'API Endpoints' section but no table.
Task ID: writer_tech_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('CloudSync Platform API Documentation', level=0)

    # --- Introduction ---
    intro = doc.add_paragraph()
    run = intro.add_run('Version 3.2.1 — Last Updated: March 2025')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(
        'This document provides a comprehensive reference for the CloudSync Platform REST API. '
        'All endpoints require authentication via Bearer token unless otherwise noted. '
        'The base URL for all API calls is https://api.cloudsync.io/v3.'
    )

    # --- Authentication Section ---
    doc.add_heading('Authentication', level=1)
    doc.add_paragraph(
        'All API requests must include a valid Bearer token in the Authorization header. '
        'Tokens can be obtained through the OAuth 2.0 flow described below.'
    )

    auth_para = doc.add_paragraph()
    run = auth_para.add_run('Authorization: Bearer <your_access_token>')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'Tokens expire after 3600 seconds. Use the refresh token endpoint to obtain a new access token '
        'without requiring the user to re-authenticate.'
    )

    # --- Rate Limiting Section ---
    doc.add_heading('Rate Limiting', level=1)
    doc.add_paragraph(
        'The API enforces rate limits to ensure fair usage across all consumers. '
        'Standard tier accounts are limited to 1000 requests per hour. '
        'Enterprise accounts receive 10,000 requests per hour. '
        'Rate limit headers are included in every response:'
    )
    doc.add_paragraph('X-RateLimit-Limit: Maximum requests per hour', style='List Bullet')
    doc.add_paragraph('X-RateLimit-Remaining: Requests remaining in current window', style='List Bullet')
    doc.add_paragraph('X-RateLimit-Reset: Unix timestamp when the limit resets', style='List Bullet')

    # --- API Endpoints Section (NO TABLE — task is to add the table) ---
    doc.add_heading('API Endpoints', level=1)
    doc.add_paragraph(
        'The following section documents all available REST API endpoints for the CloudSync Platform. '
        'Each endpoint includes the HTTP method, URL path, description, required parameters, '
        'and expected response codes.'
    )

    # --- Error Handling Section ---
    doc.add_heading('Error Handling', level=1)
    doc.add_paragraph(
        'The API uses standard HTTP status codes to indicate success or failure. '
        'Error responses include a JSON body with an error code and human-readable message.'
    )

    err_para = doc.add_paragraph()
    run = err_para.add_run('{"error": "invalid_token", "message": "The access token has expired."}')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_paragraph(
        'Common error codes include 400 (Bad Request), 401 (Unauthorized), '
        '403 (Forbidden), 404 (Not Found), and 429 (Too Many Requests).'
    )

    # --- Changelog Section ---
    doc.add_heading('Changelog', level=1)
    doc.add_paragraph('v3.2.1 (2025-03-15) — Added batch upload endpoint for file synchronization.', style='List Bullet')
    doc.add_paragraph('v3.2.0 (2025-01-20) — Introduced webhook notification support.', style='List Bullet')
    doc.add_paragraph('v3.1.0 (2024-11-05) — Added team workspace management endpoints.', style='List Bullet')
    doc.add_paragraph('v3.0.0 (2024-08-01) — Major version release with breaking changes to auth flow.', style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
