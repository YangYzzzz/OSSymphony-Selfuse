"""
Reward Script: Create a 5-column API endpoint reference table with styled header row
Task ID: writer_tech_037
Domain: libreoffice_writer
Scoring:
  Component 1: Table exists with exactly 5 columns and at least 2 rows (0.30)
  Component 2: Header row contains correct column names (0.25)
  Component 3: Header row cells have dark blue (#1565C0) background shading (0.25)
  Component 4: Header row text is white and bold (0.20)
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_037'


def color_distance_rgb(hex1, hex2):
    """Euclidean distance between two hex color strings (e.g. '1565C0')."""
    r1, g1, b1 = int(hex1[0:2], 16), int(hex1[2:4], 16), int(hex1[4:6], 16)
    r2, g2, b2 = int(hex2[0:2], 16), int(hex2[2:4], 16), int(hex2[4:6], 16)
    return sqrt((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import python-docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with exactly 5 columns and at least 2 rows (0.30 points)
    try:
        tables = doc.tables
        if len(tables) == 0:
            print(f"FAIL: Component 1 — No tables found in document")
        else:
            # Find the table near the 'API Endpoints' section
            target_table = None
            for t in tables:
                if len(t.columns) == 5:
                    target_table = t
                    break
            if target_table is None:
                # Fall back to first table
                target_table = tables[0]

            num_cols = len(target_table.columns)
            num_rows = len(target_table.rows)

            if num_cols == 5 and num_rows >= 2:
                print(f"PASS: Component 1 — Table found with {num_cols} columns and {num_rows} rows (0.30 pts)")
                total_score += 0.30
            elif num_cols == 5:
                print(f"FAIL: Component 1 — Table has 5 columns but only {num_rows} row(s), need at least 2")
            else:
                print(f"FAIL: Component 1 — Table has {num_cols} columns, expected 5")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Components 2-4 require a valid table; gate on table existence
    if len(doc.tables) == 0:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Use the 5-column table (or first table as fallback)
    target_table = None
    for t in doc.tables:
        if len(t.columns) == 5:
            target_table = t
            break
    if target_table is None:
        target_table = doc.tables[0]

    # Component 2: Header row contains correct column names (0.25 points)
    expected_headers = ['Method', 'Endpoint', 'Description', 'Parameters', 'Response Code']
    try:
        header_row = target_table.rows[0]
        actual_headers = [cell.text.strip() for cell in header_row.cells]

        # Check case-insensitive match
        matches = 0
        for exp, act in zip(expected_headers, actual_headers):
            if exp.lower() == act.lower():
                matches += 1

        if len(actual_headers) >= 5 and matches == 5:
            print(f"PASS: Component 2 — Header row has correct column names: {actual_headers} (0.25 pts)")
            total_score += 0.25
        elif matches >= 3:
            partial = round(0.25 * matches / 5, 2)
            print(f"PARTIAL: Component 2 — {matches}/5 headers match. Expected {expected_headers}, got {actual_headers} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Expected headers {expected_headers}, got {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row cells have dark blue (#1565C0) background shading (0.25 points)
    try:
        header_row = target_table.rows[0]
        cells_with_shading = 0
        TARGET_COLOR = '1565C0'
        COLOR_THRESHOLD = 40  # Euclidean RGB distance tolerance

        for ci, cell in enumerate(header_row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill.upper() != 'AUTO':
                        dist = color_distance_rgb(fill.upper(), TARGET_COLOR)
                        if dist < COLOR_THRESHOLD:
                            cells_with_shading += 1
                        else:
                            print(f"  INFO: Cell [{0},{ci}] fill={fill} distance={dist:.1f} from target {TARGET_COLOR}")

        if cells_with_shading == 5:
            print(f"PASS: Component 3 — All 5 header cells have dark blue (#1565C0) background (0.25 pts)")
            total_score += 0.25
        elif cells_with_shading >= 3:
            partial = round(0.25 * cells_with_shading / 5, 2)
            print(f"PARTIAL: Component 3 — {cells_with_shading}/5 cells have dark blue background ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {cells_with_shading}/5 header cells have dark blue background")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Header row text is white and bold (0.20 points)
    try:
        header_row = target_table.rows[0]
        white_bold_count = 0
        TARGET_WHITE = 'FFFFFF'
        WHITE_THRESHOLD = 40

        for ci, cell in enumerate(header_row.cells):
            all_runs = [run for para in cell.paragraphs for run in para.runs]
            # Check if any run has white text color (within threshold)
            has_white = any(
                run.font.color.rgb is not None
                and color_distance_rgb(str(run.font.color.rgb).upper(), TARGET_WHITE) < WHITE_THRESHOLD
                for run in all_runs
            )
            # Check if any run is bold
            has_bold = any(run.font.bold for run in all_runs)

            if has_white and has_bold:
                white_bold_count += 1

        if white_bold_count == 5:
            print(f"PASS: Component 4 — All 5 header cells have white bold text (0.20 pts)")
            total_score += 0.20
        elif white_bold_count >= 3:
            partial = round(0.20 * white_bold_count / 5, 2)
            print(f"PARTIAL: Component 4 — {white_bold_count}/5 cells have white bold text ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — Only {white_bold_count}/5 header cells have white bold text")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice state
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
