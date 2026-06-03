"""
Initial Setup: Formal letter with signature line (no underline)
Task ID: writer_txtfmt_070
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_070'
OUTPUT = f'{WORKDIR}/Desktop/formal_correspondence.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    def add_para(text, font_name='Times New Roman', font_size=12, bold=False,
                 italic=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                 space_before=0, space_after=6):
        para = doc.add_paragraph()
        para.paragraph_format.alignment = alignment
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        return para

    # Sender address block (right-aligned)
    add_para('Meridian Solutions Group', font_size=12,
             alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, space_after=0)
    add_para('1422 Harbor Boulevard, Suite 310', font_size=12,
             alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, space_after=0)
    add_para('San Francisco, CA 94105', font_size=12,
             alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, space_after=0)
    add_para('Tel: (415) 555-0178', font_size=12,
             alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT, space_after=12)

    # Date
    add_para('March 5, 2025', font_size=12, space_after=12)

    # Recipient address
    add_para('Dr. Patricia Nguyen', font_size=12, space_after=0)
    add_para('Chief Procurement Officer', font_size=12, space_after=0)
    add_para('Bellmore Industrial Holdings', font_size=12, space_after=0)
    add_para('875 Market Street, Floor 12', font_size=12, space_after=0)
    add_para('San Francisco, CA 94103', font_size=12, space_after=12)

    # Salutation
    add_para('Dear Dr. Nguyen,', font_size=12, space_after=12)

    # Body paragraphs
    add_para(
        'I am writing to follow up on our recent discussions regarding the '
        'strategic supply agreement between Meridian Solutions Group and '
        'Bellmore Industrial Holdings. We greatly appreciate the time your '
        'team has dedicated to evaluating our proposal, and we remain '
        'enthusiastic about the potential for a long-term partnership.',
        font_size=12, space_after=10
    )

    add_para(
        'Following your feedback from the meeting on February 18, we have '
        'revised the terms of our initial offer to better align with '
        'Bellmore\'s procurement cycle requirements. Specifically, we have '
        'adjusted the delivery schedule to a quarterly cadence and introduced '
        'a tiered pricing structure that reflects volume commitments over a '
        'three-year horizon.',
        font_size=12, space_after=10
    )

    add_para(
        'We believe that the revised framework addresses the key concerns '
        'raised by your logistics and finance teams. The updated contract '
        'draft has been forwarded to your legal counsel, Mr. David Harrington, '
        'for preliminary review. We anticipate that any outstanding questions '
        'can be resolved during a follow-up call scheduled for the week of '
        'March 17.',
        font_size=12, space_after=10
    )

    add_para(
        'Please do not hesitate to contact me directly should you require '
        'any additional documentation or clarification prior to that date. '
        'We look forward to concluding this agreement to the mutual benefit '
        'of both organizations.',
        font_size=12, space_after=12
    )

    # Closing
    add_para('Yours sincerely,', font_size=12, space_after=24)

    # Signature line — NO underline in initial state
    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_before = Pt(0)
    sig_para.paragraph_format.space_after = Pt(0)
    sig_run = sig_para.add_run('James T. Morrison, Director of Operations')
    sig_run.font.name = 'Times New Roman'
    sig_run.font.size = Pt(12)
    sig_run.bold = False
    sig_run.italic = False
    # Explicitly no underline
    sig_run.underline = False

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
