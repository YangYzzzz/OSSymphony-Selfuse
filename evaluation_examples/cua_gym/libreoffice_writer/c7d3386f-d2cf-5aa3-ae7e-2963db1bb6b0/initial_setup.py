"""
Initial Setup: Law thesis document with five case citations (no TOA entries marked)
Task ID: writer_acad_079
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_079'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading('Constitutional Limits on Administrative Search Powers:\nA Critical Analysis of Fourth Amendment Jurisprudence', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Jennifer L. Whitfield')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    author2 = doc.add_paragraph()
    author2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = author2.add_run('Georgetown University Law Center\nMay 2025')
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    doc.add_page_break()

    # Abstract
    doc.add_heading('Abstract', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(
        'This thesis examines the evolving interpretation of the Fourth Amendment\'s prohibition '
        'against unreasonable searches and seizures in the context of administrative regulatory '
        'inspections. Through a comprehensive analysis of landmark Supreme Court decisions, this '
        'work argues that the Court has struggled to maintain a coherent framework for balancing '
        'individual privacy rights against the government\'s legitimate regulatory interests. '
        'The thesis traces the doctrinal development from the warrant requirement established in '
        'early cases through the administrative search exception and its subsequent refinements, '
        'ultimately proposing a more principled approach to evaluating the constitutionality of '
        'warrantless administrative inspections.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_page_break()

    # Chapter I - Introduction
    doc.add_heading('Chapter I: Introduction', level=1)

    p1 = doc.add_paragraph()
    p1.paragraph_format.first_line_indent = Inches(0.5)
    p1.paragraph_format.line_spacing = 2.0
    run = p1.add_run(
        'The Fourth Amendment to the United States Constitution provides that "[t]he right of '
        'the people to be secure in their persons, houses, papers, and effects, against unreasonable '
        'searches and seizures, shall not be violated." For much of American legal history, this '
        'protection was understood primarily in the context of criminal law enforcement. However, '
        'the expansion of the modern administrative state raised fundamental questions about '
        'whether and how the Fourth Amendment constrains government regulatory inspections.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.line_spacing = 2.0
    run = p2.add_run(
        'The Supreme Court\'s decision in Camara v. Municipal Court, 387 U.S. 523 (1967) '
        'marked a watershed moment in this area of law. In Camara, the Court held for the first '
        'time that the Fourth Amendment\'s warrant requirement applies to administrative '
        'inspections of residential properties. The case arose when Roland Camara refused to '
        'permit a housing inspector to enter his residence without a warrant, and was subsequently '
        'charged with violating the San Francisco Housing Code. The Court reversed his conviction, '
        'establishing that administrative searches of private residences require either consent or '
        'a warrant, albeit one that could be issued on a lesser showing than traditional probable cause.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.paragraph_format.first_line_indent = Inches(0.5)
    p3.paragraph_format.line_spacing = 2.0
    run = p3.add_run(
        'Shortly after Camara, the Court extended this reasoning to commercial premises in '
        'See v. City of Seattle, 387 U.S. 541 (1967). In See, the Court struck down a Seattle '
        'fire code provision that authorized warrantless inspections of commercial warehouses. '
        'Justice White, writing for the majority, emphasized that "the businessman, like the '
        'occupant of a residence, has a constitutional right to go about his business free from '
        'unreasonable official entries upon his private commercial property." This decision '
        'appeared to establish a broad warrant requirement for all administrative searches.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_page_break()

    # Chapter II
    doc.add_heading('Chapter II: The Closely Regulated Industry Exception', level=1)

    p4 = doc.add_paragraph()
    p4.paragraph_format.first_line_indent = Inches(0.5)
    p4.paragraph_format.line_spacing = 2.0
    run = p4.add_run(
        'The broad warrant requirement suggested by Camara and See was soon qualified by the '
        'Court\'s recognition of the "closely regulated industry" exception. In Colonnade '
        'Catering Corp. v. United States, 397 U.S. 72 (1970), the Court addressed the '
        'constitutionality of warrantless inspections of liquor dealers. The Colonnade case '
        'involved federal agents who forcibly entered a locked storeroom at a catering establishment '
        'during a routine inspection. While the Court ultimately held that the forced entry was '
        'unauthorized by statute, it recognized the long history of government regulation of the '
        'liquor industry and suggested that Congress could authorize warrantless inspections in '
        'this context without violating the Fourth Amendment.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p5 = doc.add_paragraph()
    p5.paragraph_format.first_line_indent = Inches(0.5)
    p5.paragraph_format.line_spacing = 2.0
    run = p5.add_run(
        'The closely regulated industry doctrine reached its fullest articulation in '
        'New York v. Burger, 482 U.S. 691 (1987). In Burger, the Court upheld a New York '
        'statute authorizing warrantless inspections of automobile junkyards. Justice Blackmun, '
        'writing for the majority, set forth a three-part test for determining whether a '
        'warrantless inspection of a closely regulated business satisfies the Fourth Amendment: '
        '(1) there must be a substantial government interest that informs the regulatory scheme; '
        '(2) the warrantless inspections must be necessary to further the regulatory scheme; and '
        '(3) the inspection program must provide a constitutionally adequate substitute for a warrant.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p6 = doc.add_paragraph()
    p6.paragraph_format.first_line_indent = Inches(0.5)
    p6.paragraph_format.line_spacing = 2.0
    run = p6.add_run(
        'The Burger decision has been criticized for its expansive definition of what constitutes '
        'a "closely regulated" industry. Critics argue that virtually any business subject to '
        'licensing requirements could qualify under the Court\'s reasoning, effectively swallowing '
        'the warrant requirement established in See v. City of Seattle. Moreover, the Burger '
        'majority\'s willingness to uphold inspections that served both regulatory and law '
        'enforcement purposes raised troubling questions about the use of administrative searches '
        'as pretexts for criminal investigations.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_page_break()

    # Chapter III
    doc.add_heading('Chapter III: Modern Applications and the Digital Frontier', level=1)

    p7 = doc.add_paragraph()
    p7.paragraph_format.first_line_indent = Inches(0.5)
    p7.paragraph_format.line_spacing = 2.0
    run = p7.add_run(
        'The rise of digital technology has created new challenges for administrative search '
        'doctrine. The Supreme Court\'s landmark decision in City of Los Angeles v. Patel, '
        '576 U.S. 409 (2015) addressed whether hotel operators could be compelled to submit '
        'their guest registries for inspection without a warrant or an opportunity for '
        'precompliance review. In a 5-4 decision, the Court struck down a Los Angeles municipal '
        'code provision requiring hotels to make their records available to police upon request, '
        'holding that the Fourth Amendment requires an opportunity for precompliance review before '
        'any search of hotel records.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p8 = doc.add_paragraph()
    p8.paragraph_format.first_line_indent = Inches(0.5)
    p8.paragraph_format.line_spacing = 2.0
    run = p8.add_run(
        'Justice Sotomayor\'s majority opinion in Patel represented a significant reaffirmation '
        'of the warrant requirement in the administrative search context. The decision clarified '
        'that facial challenges to administrative search statutes are permissible under the '
        'Fourth Amendment, rejecting the argument that such challenges must be brought on an '
        'as-applied basis. Furthermore, the Court emphasized that even in closely regulated '
        'industries, some form of precompliance review is constitutionally required to prevent '
        'the government from engaging in arbitrary or harassing searches.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p9 = doc.add_paragraph()
    p9.paragraph_format.first_line_indent = Inches(0.5)
    p9.paragraph_format.line_spacing = 2.0
    run = p9.add_run(
        'The implications of Patel extend well beyond the hotel industry. As administrative '
        'agencies increasingly rely on digital records and electronic monitoring to enforce '
        'regulatory compliance, the question of what constitutional protections apply to these '
        'novel forms of surveillance becomes increasingly pressing. The Court\'s insistence on '
        'precompliance review suggests that purely automated compliance checks, without any '
        'opportunity for judicial oversight, may be constitutionally suspect.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_page_break()

    # Chapter IV - Conclusion
    doc.add_heading('Chapter IV: Conclusion', level=1)

    p10 = doc.add_paragraph()
    p10.paragraph_format.first_line_indent = Inches(0.5)
    p10.paragraph_format.line_spacing = 2.0
    run = p10.add_run(
        'This thesis has traced the development of administrative search doctrine from its '
        'origins in Camara v. Municipal Court through its modern application in City of Los '
        'Angeles v. Patel. The analysis reveals a jurisprudence marked by tension between the '
        'Court\'s commitment to Fourth Amendment protections and its recognition of the practical '
        'necessities of administrative governance. The closely regulated industry exception, as '
        'articulated in Colonnade Catering Corp. and refined in New York v. Burger, represents '
        'a significant departure from the warrant requirement, one that threatens to undermine '
        'the privacy protections that the Fourth Amendment was designed to secure.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p11 = doc.add_paragraph()
    p11.paragraph_format.first_line_indent = Inches(0.5)
    p11.paragraph_format.line_spacing = 2.0
    run = p11.add_run(
        'Looking forward, courts and legislators must grapple with the application of these '
        'doctrines to an increasingly digital and interconnected regulatory landscape. The '
        'principles established in See v. City of Seattle and reinforced in Patel provide a '
        'foundation for protecting commercial privacy in the digital age, but significant work '
        'remains to adapt these principles to new technological realities. A more principled '
        'approach would require that any warrantless administrative inspection, whether of '
        'physical premises or digital records, satisfy a heightened showing of necessity and '
        'provide meaningful procedural safeguards against arbitrary government intrusion.'
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
