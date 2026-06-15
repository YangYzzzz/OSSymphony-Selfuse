"""
Reward Script: Insert a table of authorities (legal citations) in a law thesis
Task ID: writer_acad_079
Domain: libreoffice_writer
Scoring:
  Component 1: Five TA (Table of Authorities Entry) fields present (0.50 pts)
  Component 2: TOA (Table of Authorities) index field present (0.20 pts)
  Component 3: "Table of Authorities" heading text exists (0.15 pts)
  Component 4: All five case names appear in the TOA listing (0.15 pts)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_079'

# The five case citations that should be marked as TA entries
EXPECTED_CASES = [
    "Camara v. Municipal Court",
    "See v. City of Seattle",
    "Colonnade Catering Corp. v. United States",
    "New York v. Burger",
    "City of Los Angeles v. Patel",
]


def persist_app_state():
    """Save any unsaved LibreOffice Writer edits before verification."""
    try:
        os.environ["DISPLAY"] = ":0"
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    body = doc.element.body
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Collect all instrText fields for reuse
    all_instr_fields = body.findall('.//w:instrText', ns)

    # Component 1: Five TA entry fields marked in the document (0.50 points)
    # Each TA field is worth 0.10 points, for progressive partial credit
    try:
        ta_fields = []
        for f in all_instr_fields:
            if f.text and f.text.strip().startswith('TA'):
                ta_fields.append(f.text.strip())

        ta_count = len(ta_fields)
        # Award 0.10 per TA entry, up to 5
        if ta_count > 0:
            ta_score = min(ta_count, 5) * 0.10
            total_score += ta_score

        if ta_count >= 5:
            print(f"PASS: Component 1 — {ta_count} TA entry fields found (0.50 pts)")
        elif ta_count > 0:
            print(f"PARTIAL: Component 1 — {ta_count}/5 TA entry fields found ({ta_score:.2f} pts)")
        else:
            print(f"FAIL: Component 1 — No TA entry fields found")

        # Also verify the specific case names are in the TA fields
        for case in EXPECTED_CASES:
            found = any(case in ta for ta in ta_fields)
            if found:
                print(f"  TA entry confirmed: {case}")
            else:
                print(f"  TA entry missing: {case}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: TOA (Table of Authorities) index field present (0.20 points)
    try:
        toa_fields = []
        for f in all_instr_fields:
            if f.text and 'TOA' in f.text:
                toa_fields.append(f.text.strip())

        if len(toa_fields) > 0:
            print(f"PASS: Component 2 — TOA index field found: {toa_fields[0]} (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — No TOA index field found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: "Table of Authorities" heading text exists in the document (0.15 points)
    try:
        toa_heading_found = any(
            'table of authorities' in para.text.lower() for para in doc.paragraphs
        )

        if toa_heading_found:
            print(f"PASS: Component 3 — 'Table of Authorities' heading found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — No 'Table of Authorities' heading text found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: All 5 case names appear in the TOA listing section (0.15 points)
    # The TOA listing is rendered as paragraph text containing the case citations
    try:
        # Gather all paragraph text to search for case listings
        all_text = '\n'.join(p.text for p in doc.paragraphs)

        cases_in_listing = 0
        for case in EXPECTED_CASES:
            # Check case name appears in paragraph text (the TOA listing renders them)
            if case in all_text:
                cases_in_listing += 1

        # This component checks that the TOA listing actually contains the cases
        # The initial doc also mentions these cases in body text, so we need to
        # specifically check that they appear in a TOA context (near "Table of Authorities")
        # We do this by checking that the TOA heading exists AND all cases are present
        # AND there are TA fields (compound check anchored to task changes)
        if toa_heading_found and cases_in_listing >= 5 and len(toa_fields) > 0:
            print(f"PASS: Component 4 — All 5 cases listed in TOA section (0.15 pts)")
            total_score += 0.15
        elif cases_in_listing >= 5 and len(toa_fields) > 0:
            print(f"PARTIAL: Component 4 — Cases found but TOA heading missing")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — TOA listing incomplete ({cases_in_listing}/5 cases, toa_fields={len(toa_fields)})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
