"""
Initial Setup: Internal API documentation document without watermark
Task ID: writer_tech_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with specific styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_code_block(doc, code_text):
    """Add a code-like block with monospace font."""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.5)
    return para


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # === Title Page ===
    title = doc.add_heading('Nextera Platform API', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Internal Developer Documentation')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version_para.add_run('Version 3.2.1  |  Last Updated: March 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    doc.add_paragraph()

    status_para = doc.add_paragraph()
    status_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = status_para.add_run('FOR INTERNAL USE ONLY')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # === Table of Contents ===
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '2. Authentication',
        '3. Core Endpoints',
        '   3.1 User Management',
        '   3.2 Project Resources',
        '   3.3 Analytics & Reporting',
        '4. Webhook Integration',
        '5. Error Codes & Handling',
        '6. Rate Limiting',
        '7. Versioning Policy',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # === 1. Introduction ===
    add_heading_styled(doc, '1. Introduction', level=1)
    doc.add_paragraph(
        'The Nextera Platform API provides programmatic access to all core services '
        'within the Nextera ecosystem. This RESTful API is designed to support internal '
        'tooling, third-party integrations approved by the Security team, and automated '
        'CI/CD pipelines used by the DevOps group.'
    )
    doc.add_paragraph(
        'All API communication occurs over HTTPS. The base URL for production is '
        'https://api.nextera-internal.com/v3. Staging environments use '
        'https://api-staging.nextera-internal.com/v3. All request and response bodies '
        'use JSON encoding unless otherwise noted.'
    )
    doc.add_paragraph(
        'This documentation covers the complete API surface including authentication '
        'flows, resource endpoints, webhook configuration, error handling patterns, '
        'and operational guidelines for rate limiting and versioning.'
    )

    # === 2. Authentication ===
    add_heading_styled(doc, '2. Authentication', level=1)
    doc.add_paragraph(
        'The API uses OAuth 2.0 with JWT bearer tokens for authentication. All requests '
        'must include a valid access token in the Authorization header.'
    )

    add_heading_styled(doc, '2.1 Obtaining an Access Token', level=2)
    doc.add_paragraph(
        'To obtain an access token, send a POST request to the token endpoint with your '
        'client credentials:'
    )
    add_code_block(doc,
        'POST /oauth/token\n'
        'Content-Type: application/json\n\n'
        '{\n'
        '  "grant_type": "client_credentials",\n'
        '  "client_id": "your_client_id",\n'
        '  "client_secret": "your_client_secret",\n'
        '  "scope": "read write admin"\n'
        '}'
    )
    doc.add_paragraph(
        'Tokens are valid for 3600 seconds (1 hour). The response includes a refresh_token '
        'that can be used to obtain a new access token without re-authenticating. Refresh '
        'tokens expire after 30 days of inactivity.'
    )

    add_heading_styled(doc, '2.2 Token Refresh', level=2)
    doc.add_paragraph(
        'When your access token expires, use the refresh token to obtain a new one:'
    )
    add_code_block(doc,
        'POST /oauth/token\n'
        'Content-Type: application/json\n\n'
        '{\n'
        '  "grant_type": "refresh_token",\n'
        '  "refresh_token": "your_refresh_token"\n'
        '}'
    )

    add_heading_styled(doc, '2.3 Scopes', level=2)
    doc.add_paragraph('The following permission scopes are available:')

    scope_table = doc.add_table(rows=5, cols=2)
    scope_table.style = 'Table Grid'
    scope_table.cell(0, 0).text = 'Scope'
    scope_table.cell(0, 1).text = 'Description'
    scopes = [
        ('read', 'Read-only access to all resources'),
        ('write', 'Create and update resources'),
        ('admin', 'Administrative operations including user management'),
        ('analytics', 'Access to reporting and analytics endpoints'),
    ]
    for i, (scope, desc) in enumerate(scopes, 1):
        scope_table.cell(i, 0).text = scope
        scope_table.cell(i, 1).text = desc
    # Bold header row
    for cell in scope_table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()

    # === 3. Core Endpoints ===
    add_heading_styled(doc, '3. Core Endpoints', level=1)

    # 3.1 User Management
    add_heading_styled(doc, '3.1 User Management', level=2)
    doc.add_paragraph(
        'The User Management endpoints allow you to create, retrieve, update, and '
        'deactivate user accounts within the Nextera platform.'
    )

    add_heading_styled(doc, 'GET /users', level=3)
    doc.add_paragraph(
        'Returns a paginated list of users. Supports filtering by department, role, '
        'and active status.'
    )
    add_code_block(doc,
        'GET /v3/users?department=engineering&role=developer&page=1&per_page=25\n'
        'Authorization: Bearer <access_token>\n\n'
        'Response 200:\n'
        '{\n'
        '  "data": [\n'
        '    {\n'
        '      "id": "usr_8a7b3c2d",\n'
        '      "email": "sarah.chen@nextera.com",\n'
        '      "name": "Sarah Chen",\n'
        '      "department": "Engineering",\n'
        '      "role": "Senior Developer",\n'
        '      "created_at": "2024-08-15T09:30:00Z"\n'
        '    }\n'
        '  ],\n'
        '  "pagination": { "page": 1, "per_page": 25, "total": 142 }\n'
        '}'
    )

    add_heading_styled(doc, 'POST /users', level=3)
    doc.add_paragraph(
        'Creates a new user account. Requires admin scope. An invitation email is '
        'automatically sent to the new user.'
    )
    add_code_block(doc,
        'POST /v3/users\n'
        'Authorization: Bearer <access_token>\n'
        'Content-Type: application/json\n\n'
        '{\n'
        '  "email": "marcus.johnson@nextera.com",\n'
        '  "name": "Marcus Johnson",\n'
        '  "department": "Marketing",\n'
        '  "role": "Content Manager"\n'
        '}'
    )

    # 3.2 Project Resources
    add_heading_styled(doc, '3.2 Project Resources', level=2)
    doc.add_paragraph(
        'Projects are the primary organizational unit in Nextera. Each project contains '
        'resources, team members, and configuration settings.'
    )

    add_heading_styled(doc, 'GET /projects/{project_id}', level=3)
    doc.add_paragraph('Retrieves detailed information about a specific project.')
    add_code_block(doc,
        'GET /v3/projects/prj_5f9e1a2b\n'
        'Authorization: Bearer <access_token>\n\n'
        'Response 200:\n'
        '{\n'
        '  "id": "prj_5f9e1a2b",\n'
        '  "name": "Customer Analytics Dashboard",\n'
        '  "owner": "usr_8a7b3c2d",\n'
        '  "status": "active",\n'
        '  "budget_allocated": 45230.00,\n'
        '  "team_size": 8,\n'
        '  "start_date": "2025-01-10",\n'
        '  "target_completion": "2026-06-30"\n'
        '}'
    )

    add_heading_styled(doc, 'PUT /projects/{project_id}', level=3)
    doc.add_paragraph('Updates project attributes. Only project owners and admins can modify projects.')
    add_code_block(doc,
        'PUT /v3/projects/prj_5f9e1a2b\n'
        'Authorization: Bearer <access_token>\n'
        'Content-Type: application/json\n\n'
        '{\n'
        '  "status": "on_hold",\n'
        '  "target_completion": "2026-09-30"\n'
        '}'
    )

    # 3.3 Analytics
    add_heading_styled(doc, '3.3 Analytics & Reporting', level=2)
    doc.add_paragraph(
        'The Analytics API provides access to aggregated metrics, usage statistics, '
        'and custom report generation. Data is refreshed every 15 minutes.'
    )

    add_heading_styled(doc, 'GET /analytics/summary', level=3)
    doc.add_paragraph(
        'Returns a high-level summary of platform activity for the specified date range.'
    )
    add_code_block(doc,
        'GET /v3/analytics/summary?start_date=2026-01-01&end_date=2026-03-31\n'
        'Authorization: Bearer <access_token>\n\n'
        'Response 200:\n'
        '{\n'
        '  "active_users": 1247,\n'
        '  "total_api_calls": 8923451,\n'
        '  "avg_response_time_ms": 142,\n'
        '  "error_rate": 0.0023,\n'
        '  "top_endpoints": [\n'
        '    { "path": "/users", "calls": 2341567 },\n'
        '    { "path": "/projects", "calls": 1892340 }\n'
        '  ]\n'
        '}'
    )

    # === 4. Webhook Integration ===
    add_heading_styled(doc, '4. Webhook Integration', level=1)
    doc.add_paragraph(
        'Webhooks allow your application to receive real-time notifications when events '
        'occur in the Nextera platform. You can configure webhook endpoints for specific '
        'event types through the API or the developer console.'
    )

    add_heading_styled(doc, '4.1 Registering a Webhook', level=2)
    add_code_block(doc,
        'POST /v3/webhooks\n'
        'Authorization: Bearer <access_token>\n'
        'Content-Type: application/json\n\n'
        '{\n'
        '  "url": "https://your-service.com/webhooks/nextera",\n'
        '  "events": ["user.created", "project.updated", "deployment.completed"],\n'
        '  "secret": "whsec_your_signing_secret"\n'
        '}'
    )

    doc.add_paragraph(
        'Each webhook delivery includes an X-Nextera-Signature header containing an '
        'HMAC-SHA256 signature. Always verify this signature before processing the payload '
        'to ensure the request originated from Nextera.'
    )

    # === 5. Error Codes ===
    add_heading_styled(doc, '5. Error Codes & Handling', level=1)
    doc.add_paragraph(
        'The API uses standard HTTP status codes. Error responses include a structured '
        'JSON body with details about the failure.'
    )

    error_table = doc.add_table(rows=8, cols=3)
    error_table.style = 'Table Grid'
    error_table.cell(0, 0).text = 'Code'
    error_table.cell(0, 1).text = 'Name'
    error_table.cell(0, 2).text = 'Description'
    errors = [
        ('400', 'Bad Request', 'Request body or parameters are malformed'),
        ('401', 'Unauthorized', 'Missing or invalid authentication token'),
        ('403', 'Forbidden', 'Insufficient permissions for the requested operation'),
        ('404', 'Not Found', 'The requested resource does not exist'),
        ('409', 'Conflict', 'Resource state conflict (e.g., duplicate email)'),
        ('429', 'Too Many Requests', 'Rate limit exceeded; retry after Retry-After header'),
        ('500', 'Internal Error', 'Unexpected server error; contact platform-support@nextera.com'),
    ]
    for i, (code, name, desc) in enumerate(errors, 1):
        error_table.cell(i, 0).text = code
        error_table.cell(i, 1).text = name
        error_table.cell(i, 2).text = desc
    for cell in error_table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Error response format:'
    )
    add_code_block(doc,
        '{\n'
        '  "error": {\n'
        '    "code": "VALIDATION_ERROR",\n'
        '    "message": "The email field is required.",\n'
        '    "details": [\n'
        '      { "field": "email", "issue": "missing_required_field" }\n'
        '    ],\n'
        '    "request_id": "req_7c4d2e1f"\n'
        '  }\n'
        '}'
    )

    # === 6. Rate Limiting ===
    add_heading_styled(doc, '6. Rate Limiting', level=1)
    doc.add_paragraph(
        'To ensure fair usage and platform stability, the API enforces rate limits on '
        'all endpoints. Rate limits are applied per API key.'
    )

    rate_table = doc.add_table(rows=4, cols=3)
    rate_table.style = 'Table Grid'
    rate_table.cell(0, 0).text = 'Tier'
    rate_table.cell(0, 1).text = 'Requests/Minute'
    rate_table.cell(0, 2).text = 'Burst Limit'
    rates = [
        ('Standard', '60', '100'),
        ('Professional', '300', '500'),
        ('Enterprise', '1000', '2000'),
    ]
    for i, (tier, rpm, burst) in enumerate(rates, 1):
        rate_table.cell(i, 0).text = tier
        rate_table.cell(i, 1).text = rpm
        rate_table.cell(i, 2).text = burst
    for cell in rate_table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Rate limit status is returned in response headers: X-RateLimit-Limit, '
        'X-RateLimit-Remaining, and X-RateLimit-Reset. When the limit is exceeded, '
        'the API returns HTTP 429 with a Retry-After header indicating seconds until '
        'the next available request window.'
    )

    # === 7. Versioning ===
    add_heading_styled(doc, '7. Versioning Policy', level=1)
    doc.add_paragraph(
        'The API follows semantic versioning. The current version is v3. Breaking changes '
        'are only introduced in major version increments. Deprecated endpoints receive a '
        'minimum 6-month sunset period with advance notice via the developer changelog.'
    )
    doc.add_paragraph(
        'You can specify the API version in the URL path (/v3/...) or via the '
        'X-API-Version request header. If no version is specified, the latest stable '
        'version is used. We strongly recommend pinning to a specific version in '
        'production integrations.'
    )
    doc.add_paragraph(
        'For questions or support, contact the Platform Engineering team at '
        'platform-support@nextera.com or via the #api-support Slack channel.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
