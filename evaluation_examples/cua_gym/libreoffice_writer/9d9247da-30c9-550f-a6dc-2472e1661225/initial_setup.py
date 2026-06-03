"""
Initial Setup: Create design_review.docx with dense technical text (page 1) and blueprint.png
Task ID: writer_obj_069
Domain: libreoffice_writer

No inserted image in the docx - agent needs to insert blueprint.png from Desktop.
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_069'
DESKTOP = f'{WORKDIR}/Desktop'
DOC_OUTPUT = f'{DESKTOP}/design_review.docx'
IMG_OUTPUT = f'{DESKTOP}/blueprint.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_blueprint_png():
    """Create a 1000x750 PNG image with transparent areas around main content."""
    from PIL import Image, ImageDraw, ImageFont

    # Create RGBA image with transparency
    img = Image.new('RGBA', (1000, 750), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    # Draw blueprint-style content in the center (non-transparent area)
    # Background fill for main content area
    draw.rectangle([50, 50, 950, 700], fill=(20, 40, 80, 240))

    # Grid lines
    for x in range(50, 951, 50):
        draw.line([(x, 50), (x, 700)], fill=(60, 100, 160, 180), width=1)
    for y in range(50, 701, 50):
        draw.line([(50, y), (950, y)], fill=(60, 100, 160, 180), width=1)

    # Main structural lines (blueprint style)
    draw.rectangle([100, 100, 900, 650], outline=(150, 200, 255, 255), width=3)
    draw.rectangle([150, 150, 850, 600], outline=(100, 180, 255, 255), width=2)

    # Floor plan-like elements
    # Room 1 (left)
    draw.rectangle([160, 160, 450, 380], outline=(200, 220, 255, 255), width=2)
    # Room 2 (right)
    draw.rectangle([500, 160, 840, 380], outline=(200, 220, 255, 255), width=2)
    # Room 3 (bottom-left)
    draw.rectangle([160, 420, 380, 590], outline=(200, 220, 255, 255), width=2)
    # Room 4 (bottom-right)
    draw.rectangle([430, 420, 840, 590], outline=(200, 220, 255, 255), width=2)

    # Dimension lines
    draw.line([(160, 630), (840, 630)], fill=(255, 220, 50, 220), width=2)
    draw.line([(160, 625), (160, 635)], fill=(255, 220, 50, 220), width=2)
    draw.line([(840, 625), (840, 635)], fill=(255, 220, 50, 220), width=2)

    # Labels using basic drawing (no font file needed)
    # Title bar
    draw.rectangle([100, 660, 900, 690], fill=(30, 60, 120, 220))

    # Border (transparent corners to simulate contour)
    # Make corners more transparent
    for corner_x, corner_y in [(0, 0), (950, 0), (0, 700), (950, 700)]:
        for dx in range(50):
            for dy in range(50):
                dist = ((dx)**2 + (dy)**2) ** 0.5
                if dist < 50:
                    alpha = int(dist / 50 * 100)
                    px = corner_x + dx if corner_x == 0 else corner_x - dx
                    py = corner_y + dy if corner_y == 0 else corner_y - dy
                    if 0 <= px < 1000 and 0 <= py < 750:
                        r, g, b, a = img.getpixel((px, py))
                        img.putpixel((px, py), (r, g, b, min(a, alpha)))

    img.save(IMG_OUTPUT, 'PNG')
    print(f'Blueprint image created: {IMG_OUTPUT}')


def create_design_review_docx():
    """Create design_review.docx with dense technical text on page 1. No inserted blueprint image."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Title
    title = doc.add_heading('Technical Design Review — Project Helios', level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x7A)

    # Subtitle / metadata
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    run = meta.add_run('Document Reference: HDR-2025-003   |   Classification: INTERNAL   |   Rev: 2.1')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.italic = True

    # Section 1
    h2 = doc.add_heading('1. Executive Summary', level=2)
    h2.runs[0].font.size = Pt(13)

    p1 = doc.add_paragraph(
        'This document presents a comprehensive technical design review for the structural reconfiguration '
        'of Building Complex C at the Meridian Industrial Campus. The review encompasses load-bearing '
        'analysis, HVAC system integration, electrical distribution mapping, and safety egress planning '
        'as required under ISO 45001 and local building code EN-1992-1-1. All measurements conform to '
        'metric SI units; tolerances are expressed as ±0.5 mm unless otherwise stated.'
    )
    p1.paragraph_format.space_after = Pt(6)
    p1.runs[0].font.size = Pt(10.5)

    p2 = doc.add_paragraph(
        'The primary objective of Phase II is to expand the eastern wing by 2,400 m² while maintaining '
        'full operational continuity of the existing manufacturing floor. Structural analysis indicates '
        'that the proposed 12-column grid extension (6 m × 8 m bay spacing) is feasible with reinforced '
        'concrete footings rated at 450 kN/m². Dynamic load calculations include wind exposure Category B '
        'and seismic zone classification SZ-IIa.'
    )
    p2.paragraph_format.space_after = Pt(6)
    p2.runs[0].font.size = Pt(10.5)

    # Section 2
    h2b = doc.add_heading('2. Structural Load Analysis', level=2)
    h2b.runs[0].font.size = Pt(13)

    p3 = doc.add_paragraph(
        'Dead loads for the proposed extension are calculated at 5.2 kN/m² for the composite deck and '
        '3.8 kN/m² for the cladding assembly. Live loads are specified at 7.5 kN/m² for the production '
        'floor and 3.0 kN/m² for the mezzanine storage level. Wind pressure coefficients have been '
        'derived from CFD simulations performed at Reynolds number Re = 2.4 × 10⁷, yielding a design '
        'pressure of 1.2 kPa on the windward facade.'
    )
    p3.paragraph_format.space_after = Pt(6)
    p3.runs[0].font.size = Pt(10.5)

    p4 = doc.add_paragraph(
        'Foundation design utilises bored piles of 600 mm diameter at 3 m centres, achieving a nominal '
        'bearing capacity of 2,200 kN per pile. Settlement analysis using the Terzaghi consolidation '
        'model predicts a maximum long-term differential settlement of 18 mm at the expansion joint '
        'between the existing and new structures — within the permissible 25 mm limit. Pile reinforcement '
        'consists of 8T25 longitudinal bars with R10 links at 200 mm c/c throughout the pile shaft.'
    )
    p4.paragraph_format.space_after = Pt(6)
    p4.runs[0].font.size = Pt(10.5)

    # Section 3
    h2c = doc.add_heading('3. Mechanical & Electrical Systems', level=2)
    h2c.runs[0].font.size = Pt(13)

    p5 = doc.add_paragraph(
        'The HVAC system expansion incorporates four Daikin VRV-IV heat recovery units (nominal capacity '
        '112 kW each) connected to a variable-flow refrigerant network. Supply air is delivered at '
        '18°C ± 1°C via 600 mm × 400 mm ductwork running at a maximum face velocity of 6 m/s. '
        'Exhaust air extraction is achieved through 12 exhaust points each rated at 2,500 m³/h, '
        'maintaining a positive building pressure of +15 Pa relative to ambient.'
    )
    p5.paragraph_format.space_after = Pt(6)
    p5.runs[0].font.size = Pt(10.5)

    p6 = doc.add_paragraph(
        'Power distribution for the extended wing is fed from a new 1,600 kVA dry-type transformer '
        '(11 kV / 415 V, Dyn11 vector group) installed in substation room SS-07. Main LV switchboard '
        'MLVS-07 is equipped with 2,000 A ACB incoming, 16 × 250 A MCCBs for distribution circuits, '
        'and 4 × 100 A MCBs for lighting sub-circuits. Cable routes follow BIM-coordinated cable ladder '
        'systems per IEC 61537 with minimum bending radii observed throughout.'
    )
    p6.paragraph_format.space_after = Pt(6)
    p6.runs[0].font.size = Pt(10.5)

    # Section 4
    h2d = doc.add_heading('4. Fire Safety & Egress Planning', level=2)
    h2d.runs[0].font.size = Pt(13)

    p7 = doc.add_paragraph(
        'Fire compartmentation is achieved through 2-hour fire-rated walls (90 mm blockwork with '
        '13 mm plasterboard lining, U-value 0.45 W/m²K) and intumescent fire stopping at all service '
        'penetrations. Emergency egress routes comply with EN 1838 illumination levels (1 lux minimum '
        'on escape route floor) with luminaire spacing of 8 m maximum. Exit signage uses photoluminescent '
        'material to EN ISO 7010-E001 standard, maintaining visibility for 60 minutes post-power failure.'
    )
    p7.paragraph_format.space_after = Pt(6)
    p7.runs[0].font.size = Pt(10.5)

    p8 = doc.add_paragraph(
        'Sprinkler coverage follows NFPA 13 Ordinary Hazard Group 2 classification with K-80 upright '
        'sprinkler heads at 3.0 m × 3.0 m grid spacing. Design density is 8.1 mm/min over a design '
        'area of 260 m². Fire pump assembly comprises a diesel-driven main pump (750 L/min @ 7 bar) '
        'with an electric jockey pump (120 L/min @ 8 bar) and 45 m³ storage tank providing 60 minutes '
        'reserve capacity.'
    )
    p8.paragraph_format.space_after = Pt(6)
    p8.runs[0].font.size = Pt(10.5)

    # Section 5 — starts on same page, continues later
    h2e = doc.add_heading('5. Drawing Reference Index', level=2)
    h2e.runs[0].font.size = Pt(13)

    p9 = doc.add_paragraph(
        'The following drawings constitute the primary contract documentation set. All drawings are '
        'produced in AutoCAD 2025 and exported at 1:100 scale unless noted. Revisions are tracked '
        'via the Document Management System (DMS) under project code HLV-C-2025.'
    )
    p9.paragraph_format.space_after = Pt(6)
    p9.runs[0].font.size = Pt(10.5)

    # Table of drawing references
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(['Drawing No.', 'Title', 'Scale', 'Rev.']):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    drawing_data = [
        ('C-001', 'Site Layout & Boundary Survey', '1:500', 'B'),
        ('C-002', 'Ground Floor Plan — Existing', '1:100', 'C'),
        ('C-003', 'Ground Floor Plan — Proposed Extension', '1:100', 'A'),
        ('C-004', 'First Floor Mezzanine Layout', '1:100', 'A'),
        ('C-005', 'Roof Plan & Drainage Layout', '1:100', 'B'),
        ('S-001', 'Foundation Layout & Pile Schedule', '1:100', 'A'),
        ('S-002', 'Column & Beam Grid Lines', '1:100', 'A'),
        ('S-003', 'Structural Sections A-A & B-B', '1:50', 'A'),
        ('M-001', 'HVAC Ductwork Ground Floor', '1:100', 'A'),
        ('M-002', 'Refrigerant Piping Schematic', 'NTS', 'A'),
        ('E-001', 'LV Distribution Single Line Diagram', 'NTS', 'A'),
        ('E-002', 'Cable Routing Ground Floor', '1:100', 'A'),
        ('F-001', 'Fire Compartmentation Plan', '1:100', 'A'),
        ('F-002', 'Sprinkler Layout Ground Floor', '1:100', 'A'),
    ]
    for row_data in drawing_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            run = row.cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph('')  # spacing

    p10 = doc.add_paragraph(
        'All referenced drawings shall be read in conjunction with the Structural Engineer\'s '
        'specification (document ref. SE-SPEC-2025-HLV) and the M&E Coordination Drawing Register '
        '(document ref. MEP-CDR-2025-003). Any discrepancies between drawings and specifications '
        'shall be immediately notified to the project lead engineer, Dr. Amara Singh (ext. 4412), '
        'for resolution prior to commencement of affected works.'
    )
    p10.paragraph_format.space_after = Pt(6)
    p10.runs[0].font.size = Pt(10.5)

    os.makedirs(DESKTOP, exist_ok=True)
    doc.save(DOC_OUTPUT)
    print(f'Initial document created: {DOC_OUTPUT}')


def create_initial():
    create_blueprint_png()
    create_design_review_docx()

    # GUI-ready startup: open the docx in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with design_review.docx (DISPLAY=:0)')


create_initial()
