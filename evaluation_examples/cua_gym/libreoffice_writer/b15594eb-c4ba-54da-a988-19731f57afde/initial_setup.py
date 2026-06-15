"""
Initial Setup: Budget Plan expense table with empty Total column
Task ID: writer_af_017
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_af_017'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    heading = doc.add_heading("Q2 2025 Budget Plan", level=1)

    # Intro paragraph
    doc.add_paragraph(
        "The following table outlines the projected expenses for the "
        "second quarter of 2025. Please review each line item and verify "
        "the totals before submission to the finance department."
    )

    # Expense table: 8 rows (1 header + 6 data + 1 grand total), 4 columns
    table = doc.add_table(rows=8, cols=4)
    table.style = "Table Grid"

    # Header row
    headers = ["Item", "Quantity", "Unit Price", "Total"]
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)

    # Data rows (6 rows of realistic expense items)
    data = [
        ["Office Supplies (Paper, Pens, Toner)", 25, 18.50],
        ["Ergonomic Desk Chairs", 8, 349.99],
        ["Software Licenses (Annual)", 15, 129.00],
        ["Conference Room AV Equipment", 3, 1275.00],
        ["Employee Training Materials", 12, 85.75],
        ["Cloud Hosting Services (Monthly)", 6, 450.00],
    ]

    for row_idx, (item, qty, price) in enumerate(data, start=1):
        table.cell(row_idx, 0).text = item
        table.cell(row_idx, 1).text = str(qty)
        table.cell(row_idx, 2).text = f"{price:.2f}"
        # Column D (Total) is intentionally left EMPTY
        table.cell(row_idx, 3).text = ""

    # Grand total row - label in column A, column D empty
    grand_total_cell_a = table.cell(7, 0)
    grand_total_cell_a.text = ""
    run_gt = grand_total_cell_a.paragraphs[0].add_run("Grand Total")
    run_gt.bold = True
    run_gt.font.size = Pt(11)
    table.cell(7, 1).text = ""
    table.cell(7, 2).text = ""
    table.cell(7, 3).text = ""  # Grand total cell is EMPTY

    # Additional context paragraph
    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: All prices are listed in USD. The Total column should "
        "reflect Quantity multiplied by Unit Price for each item. "
        "The Grand Total should be the sum of all item totals."
    )

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
