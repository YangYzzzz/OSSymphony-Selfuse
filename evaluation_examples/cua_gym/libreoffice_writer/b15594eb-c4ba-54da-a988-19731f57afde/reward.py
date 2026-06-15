"""
Reward Script: Add formulas to expense table in Writer document
Task ID: writer_af_017
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Data rows column D have correct Quantity * Unit Price values
  Component 2 (0.3): Grand total row has correct sum of column D
  Component 3 (0.2): All column D values are numeric and non-empty
"""

import os
import time


WORKDIR = '/home/user'
TASK_ID = 'writer_af_017'

# Expected values derived from task data:
# Row 1: 25 * 18.50 = 462.50
# Row 2: 8 * 349.99 = 2799.92
# Row 3: 15 * 129.00 = 1935.00
# Row 4: 3 * 1275.00 = 3825.00
# Row 5: 12 * 85.75 = 1029.00
# Row 6: 6 * 450.00 = 2700.00
# Grand Total: 12751.42
EXPECTED_ROW_VALUES = {
    1: 462.50,
    2: 2799.92,
    3: 1935.00,
    4: 3825.00,
    5: 1029.00,
    6: 2700.00,
}
EXPECTED_GRAND_TOTAL = 12751.42


def persist_app_state(domain: str):
    """Attempt to save any unsaved LibreOffice state via Ctrl+S."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("CRITICAL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Precondition: table must have 8 rows and 4 columns
    if len(table.rows) < 8 or len(table.columns) < 4:
        print(f"CRITICAL: Table dimensions unexpected: {len(table.rows)} rows x {len(table.columns)} cols")
        print("REWARD: 0.0")
        return 0.0

    # Helper: parse a cell value as a float
    def parse_float(text):
        """Try to parse cell text as a float. Returns None on failure."""
        text = text.strip().replace(',', '').replace('$', '').replace(' ', '')
        if not text:
            return None
        try:
            return float(text)
        except ValueError:
            return None

    # Component 1: Data rows (1-6) in column D have correct Quantity * Unit Price (0.5 points)
    # Each correct row earns ~0.0833 points (0.5 / 6)
    try:
        correct_rows = 0
        per_row_score = 0.5 / 6.0
        for row_idx in range(1, 7):
            cell_d_text = table.rows[row_idx].cells[3].text.strip()
            actual_val = parse_float(cell_d_text)
            expected_val = EXPECTED_ROW_VALUES[row_idx]
            if actual_val is not None and abs(actual_val - expected_val) < 0.02:
                correct_rows += 1
                print(f"PASS: Row {row_idx} Col D = {actual_val} (expected {expected_val})")
            else:
                print(f"FAIL: Row {row_idx} Col D = {repr(cell_d_text)} (expected {expected_val})")
        row_score = correct_rows * per_row_score
        if correct_rows > 0:
            print(f"PASS: Component 1 — {correct_rows}/6 data rows correct ({row_score:.3f} pts)")
            total_score += row_score
        else:
            print(f"FAIL: Component 1 — no data rows have correct values")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Grand total row (row 7, col D) has the correct sum (0.3 points)
    try:
        grand_total_text = table.rows[7].cells[3].text.strip()
        grand_total_val = parse_float(grand_total_text)
        if grand_total_val is not None and abs(grand_total_val - EXPECTED_GRAND_TOTAL) < 0.02:
            print(f"PASS: Component 2 — Grand total = {grand_total_val} (expected {EXPECTED_GRAND_TOTAL}) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Grand total = {repr(grand_total_text)} (expected {EXPECTED_GRAND_TOTAL})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All 7 column D values (rows 1-7) are numeric and non-empty (0.2 points)
    # This checks that the agent filled ALL cells, not just some
    try:
        all_numeric = 0
        for row_idx in range(1, 8):
            cell_text = table.rows[row_idx].cells[3].text.strip()
            val = parse_float(cell_text)
            if val is not None and val > 0:
                all_numeric += 1
        if all_numeric == 7:
            print(f"PASS: Component 3 — All 7 column D cells are numeric and non-empty (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Only {all_numeric}/7 column D cells are numeric ({[table.rows[r].cells[3].text.strip() for r in range(1,8)]})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice state
persist_app_state("libreoffice_writer")

# Run verification
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
