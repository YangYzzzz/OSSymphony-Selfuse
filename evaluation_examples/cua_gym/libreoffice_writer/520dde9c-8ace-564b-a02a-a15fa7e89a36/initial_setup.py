"""
Initial Setup: Recipe instructions document with steps running together in first paragraph
Task ID: osworld_writer_spacing_010
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_spacing_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("Grandma's Classic Lasagna", level=1)

    # First paragraph: 8 cooking steps all running together as one block
    # Each sentence is a step ending with a period — the agent must separate them
    steps_block = (
        "Preheat the oven to 375 degrees Fahrenheit. "
        "Grease a 9x13 baking pan thoroughly. "
        "Mix the dry ingredients in a large bowl. "
        "Brown the ground beef with onion and garlic in a skillet. "
        "Simmer the meat sauce for 20 minutes over medium heat. "
        "Layer noodles, ricotta mixture, meat sauce, and mozzarella cheese. "
        "Repeat layers until all ingredients are used, finishing with cheese. "
        "Bake covered with foil for 45 minutes, then uncovered for 15 minutes."
    )
    first_para = doc.add_paragraph(steps_block)
    first_para.paragraph_format.space_before = Pt(0)
    first_para.paragraph_format.space_after = Pt(0)

    # Second paragraph: serving suggestions — must remain untouched
    serving = doc.add_paragraph(
        "Serving suggestions: Let the lasagna rest for 10 minutes before slicing. "
        "Serve with a side of garlic bread and a fresh green salad. "
        "Pairs well with a glass of Chianti or sparkling water with lemon. "
        "Leftovers keep well in the refrigerator for up to 4 days."
    )
    serving.paragraph_format.space_before = Pt(12)
    serving.paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
