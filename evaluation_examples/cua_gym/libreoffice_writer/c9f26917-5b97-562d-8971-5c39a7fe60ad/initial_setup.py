"""
Initial Setup: Master document with subdocuments where all sections restart page numbering.
Task ID: writer_rm_062
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_062'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph, num_format='decimal'):
    """Add a PAGE field code to a paragraph with specified number format.
    num_format: 'decimal' for 1,2,3 or 'lowerRoman' for i,ii,iii
    """
    from docx.oxml.ns import qn

    # Begin field char
    r1 = paragraph.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    # Instruction text
    r2 = paragraph.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    # End field char
    r3 = paragraph.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def set_page_number_format(section, fmt='decimal', start=None):
    """Set page number format on a section.
    fmt: 'decimal', 'lowerRoman', 'upperRoman', 'lowerLetter', 'upperLetter'
    start: page number to start at (integer), or None
    """
    from docx.oxml.ns import qn

    sectPr = section._sectPr
    # Remove existing pgNumType if present
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)

    attrs = {qn('w:fmt'): fmt}
    if start is not None:
        attrs[qn('w:start')] = str(start)
    pgNumType = sectPr.makeelement(qn('w:pgNumType'), attrs)
    sectPr.append(pgNumType)


def set_section_start(section, start_type='nextPage'):
    """Set section start type via XML."""
    from docx.oxml.ns import qn

    sectPr = section._sectPr
    # Remove existing type if present
    for existing in sectPr.findall(qn('w:type')):
        sectPr.remove(existing)

    type_elem = sectPr.makeelement(qn('w:type'), {qn('w:val'): start_type})
    sectPr.append(type_elem)


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn

    doc = Document()

    # ---- Default section = Preface ----
    section0 = doc.sections[0]
    section0.page_width = Inches(8.5)
    section0.page_height = Inches(11)
    section0.left_margin = Inches(1)
    section0.right_margin = Inches(1)
    section0.top_margin = Inches(1)
    section0.bottom_margin = Inches(1)

    # Preface title
    h = doc.add_heading('Preface', level=1)
    h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    preface_text = [
        "This book represents a culmination of five years of research into distributed "
        "computing architectures and their applications in modern enterprise systems. What "
        "began as a series of lectures at Stanford University has evolved into a comprehensive "
        "guide for practitioners and researchers alike.",

        "The field of distributed systems has undergone remarkable transformation since the "
        "publication of the seminal works by Lamport and others in the 1970s. Today, with the "
        "proliferation of cloud computing platforms and microservices architectures, the "
        "principles discussed in these chapters are more relevant than ever.",

        "I wish to extend my gratitude to the many colleagues who reviewed early drafts of "
        "this manuscript: Dr. Elena Rodriguez at MIT, Professor Hiroshi Tanaka at the "
        "University of Tokyo, and the entire distributed systems team at Meridian Technologies. "
        "Their insights have been invaluable.",

        "Special thanks go to my editor, Margaret Thornton, whose patience and keen eye for "
        "clarity transformed dense technical prose into accessible content. Any remaining "
        "errors are, of course, entirely my own.",
    ]
    for text in preface_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    # Preface footer with Roman numeral page number
    footer0 = section0.footer
    footer0.is_linked_to_previous = False
    fp0 = footer0.paragraphs[0]
    fp0.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fp0.add_run("- ")
    add_page_number_field(fp0, num_format='lowerRoman')
    fp0.add_run(" -")

    # Set preface page numbering: lower Roman, starting at i
    set_page_number_format(section0, fmt='lowerRoman', start=1)

    # ---- Chapters 1-5 ----
    chapter_titles = [
        "Chapter 1: Foundations of Distributed Computing",
        "Chapter 2: Consensus Algorithms and Fault Tolerance",
        "Chapter 3: Data Replication Strategies",
        "Chapter 4: Scalable Message Passing Systems",
        "Chapter 5: Case Studies in Production Systems",
    ]

    chapter_contents = [
        [
            "Distributed computing begins with a fundamental premise: that computation can be "
            "divided across multiple processing units to achieve goals that no single machine "
            "could accomplish efficiently. This chapter establishes the theoretical framework "
            "upon which the rest of the book is built.",

            "We start by examining the CAP theorem, first conjectured by Eric Brewer in 2000 "
            "and formally proved by Gilbert and Lynch in 2002. The theorem states that a "
            "distributed data store cannot simultaneously provide more than two of the following "
            "three guarantees: Consistency, Availability, and Partition tolerance.",

            "Consider a simple banking application deployed across three data centers in "
            "New York, London, and Singapore. When a customer in London initiates a transfer "
            "of $10,000, the system must decide how to handle the request if the network "
            "connection to the New York data center is temporarily severed.",

            "The FLP impossibility result, published by Fischer, Lynch, and Paterson in 1985, "
            "provides another fundamental constraint. It demonstrates that in an asynchronous "
            "distributed system, no deterministic consensus algorithm can guarantee agreement "
            "in the presence of even a single process failure.",
        ],
        [
            "Consensus lies at the heart of distributed systems. When multiple nodes must "
            "agree on a single value or sequence of operations, the challenge of achieving "
            "reliable consensus becomes paramount. This chapter examines the major consensus "
            "protocols that have shaped the field.",

            "The Paxos algorithm, introduced by Leslie Lamport in 1998, was the first "
            "practical solution to consensus in asynchronous networks with crash failures. "
            "Despite its theoretical elegance, Paxos proved notoriously difficult to implement "
            "correctly, leading to years of engineering effort at companies like Google.",

            "Raft, proposed by Ongaro and Ousterhout in 2014, was designed explicitly for "
            "understandability. It decomposes the consensus problem into three relatively "
            "independent subproblems: leader election, log replication, and safety. The "
            "algorithm has been adopted by systems such as etcd, CockroachDB, and TiKV.",

            "Byzantine fault tolerance (BFT) addresses a more challenging scenario where "
            "nodes may behave arbitrarily, including sending conflicting information to "
            "different peers. The Practical BFT (PBFT) algorithm by Castro and Liskov "
            "requires 3f + 1 nodes to tolerate f Byzantine faults.",
        ],
        [
            "Data replication is essential for both fault tolerance and performance in "
            "distributed systems. By maintaining copies of data across multiple nodes, "
            "systems can continue operating even when individual components fail, while "
            "also serving read requests from geographically closer replicas.",

            "Synchronous replication ensures that all replicas are updated before a write "
            "is acknowledged to the client. While this provides strong consistency guarantees, "
            "it comes at the cost of increased latency and reduced availability. PostgreSQL's "
            "synchronous replication mode exemplifies this approach.",

            "Asynchronous replication, by contrast, acknowledges writes immediately and "
            "propagates changes to replicas in the background. Amazon DynamoDB and Apache "
            "Cassandra employ variations of this approach, trading consistency for lower "
            "latency and higher availability.",

            "Conflict-free Replicated Data Types (CRDTs) represent an innovative approach "
            "to the replication problem. By designing data structures whose operations are "
            "commutative, CRDTs guarantee that all replicas eventually converge to the same "
            "state without requiring coordination.",
        ],
        [
            "Message passing forms the backbone of communication in distributed systems. "
            "Whether using synchronous RPC calls or asynchronous message queues, the choice "
            "of communication pattern profoundly affects system behavior, performance, and "
            "reliability.",

            "Apache Kafka has emerged as the de facto standard for high-throughput event "
            "streaming. Its log-based architecture, combined with consumer group semantics, "
            "enables both real-time stream processing and batch analytics on the same data. "
            "LinkedIn processes over 7 trillion messages per day through Kafka.",

            "gRPC, developed by Google, provides a modern framework for synchronous "
            "communication between services. Built on HTTP/2, it supports bidirectional "
            "streaming, flow control, and efficient binary serialization through Protocol "
            "Buffers. Adoption has grown rapidly across the industry.",

            "The actor model, implemented in frameworks like Akka and Microsoft Orleans, "
            "provides an alternative paradigm where lightweight actors communicate exclusively "
            "through message passing. This eliminates shared mutable state, greatly simplifying "
            "reasoning about concurrent and distributed computations.",
        ],
        [
            "Theory and practice often diverge in distributed systems. This chapter presents "
            "detailed case studies from production systems at major technology companies, "
            "illustrating how theoretical concepts are adapted to real-world constraints.",

            "Google Spanner represents a landmark achievement in distributed databases. By "
            "using GPS-synchronized atomic clocks (TrueTime), Spanner provides externally "
            "consistent reads and writes across globally distributed data centers. The system "
            "supports the full SQL feature set while maintaining strong consistency.",

            "Amazon Aurora reimagines the traditional database architecture by separating "
            "compute from storage. The storage layer uses a distributed, log-structured "
            "design with six-way replication across three availability zones, achieving "
            "both high availability and strong durability guarantees.",

            "Netflix's microservices architecture, built on AWS, serves over 230 million "
            "subscribers worldwide. Their approach to resilience engineering, including "
            "the Chaos Monkey tool and circuit breaker patterns, has influenced the entire "
            "industry's thinking about building reliable distributed systems.",
        ],
    ]

    for i in range(5):
        # Add section break for each chapter
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        section = doc.sections[i + 1]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # Chapter heading
        h = doc.add_heading(chapter_titles[i], level=1)
        h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Chapter content
        for text in chapter_contents[i]:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

        # Footer with page number (Arabic) - each chapter restarts
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fp.add_run("- ")
        add_page_number_field(fp, num_format='decimal')
        fp.add_run(" -")

        # INITIAL STATE: All chapters RESTART page numbering at 1
        # This is the problem the agent needs to fix
        set_page_number_format(section, fmt='decimal', start=1)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
