"""
Initial Setup: 6-page magazine-style article in single-column layout
Task ID: writer_page_047
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_page_047'
OUTPUT = f'{WORKDIR}/Desktop/magazine_layout.docx'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_single_column(section):
    """Configure section to use a single column layout."""
    sectPr = section._sectPr
    # Remove any existing cols elements
    for c in sectPr.findall(f'{{{W_NS}}}cols'):
        sectPr.remove(c)
    # Add single-column specification
    cols = etree.SubElement(sectPr, f'{{{W_NS}}}cols')
    cols.set(f'{{{W_NS}}}num', '1')
    cols.set(f'{{{W_NS}}}space', '708')  # default 1.25cm spacing (unused for single col)


def set_a4_portrait(section):
    """Set A4 portrait page with margins top=1.5cm, bottom=1.5cm, left=1.5cm, right=1.5cm."""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def create_initial():
    # Ensure Desktop directory exists on VM
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()
    section = doc.sections[0]

    # Set A4 portrait with specified margins
    set_a4_portrait(section)

    # Set single column layout (MUST NOT be 3-column before task)
    set_single_column(section)

    # ---- Page 1: Cover / Introduction ----
    heading = doc.add_heading('The Future of Urban Living', level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph('A Special Report on Modern City Design & Architecture')
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sub.runs[0]
    run.italic = True
    run.font.size = Pt(13)

    doc.add_paragraph()  # spacer

    intro = doc.add_paragraph(
        'As cities around the world grapple with rapid urbanization, climate change, '
        'and shifting demographics, urban planners and architects are reimagining '
        'what the modern city can look like. From vertical gardens to smart infrastructure, '
        'a new generation of solutions is transforming the urban landscape.'
    )
    intro.paragraph_format.space_after = Pt(10)

    doc.add_paragraph(
        'This magazine explores the most innovative projects currently reshaping '
        'our cities — examining how technology, sustainability, and community-centered '
        'design are coming together to create more livable urban environments for everyone.'
    )

    doc.add_page_break()

    # ---- Page 2: Feature Article 1 ----
    doc.add_heading('Green Rooftops: The Sky\'s the Limit', level=1)

    doc.add_paragraph(
        'Copenhagen, Singapore, and New York City have emerged as global leaders in '
        'rooftop agriculture and green infrastructure. By converting previously unused '
        'rooftop space into lush gardens, urban farms, and biodiversity havens, cities '
        'are fighting the urban heat island effect while providing fresh produce to '
        'local communities.'
    )

    doc.add_paragraph(
        'The Bosco Verticale in Milan stands as one of the most celebrated examples '
        'of vertical urban greening. Designed by Stefano Boeri Architetti, the twin '
        'residential towers house over 800 trees and 15,000 plants on their terraces '
        'and balconies — the equivalent of a small forest suspended in the sky.'
    )

    doc.add_heading('Key Benefits of Urban Greening', level=2)

    benefits = [
        'Reduction of urban heat island effect by up to 3°C',
        'Improved air quality through natural filtration',
        'Stormwater management and reduced urban flooding',
        'Increased biodiversity and wildlife corridors',
        'Mental health benefits for urban residents',
        'Energy savings through natural insulation',
    ]
    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_page_break()

    # ---- Page 3: Feature Article 2 ----
    doc.add_heading('Smart Mobility: Beyond the Private Car', level=1)

    doc.add_paragraph(
        'The internal combustion engine dominated urban mobility for over a century. '
        'But as cities face gridlock, pollution, and the existential threat of climate '
        'change, a fundamental rethinking of how people move through urban spaces is '
        'underway. The results are nothing short of revolutionary.'
    )

    doc.add_paragraph(
        'Amsterdam\'s cycling infrastructure has long served as a model for car-free '
        'urban planning. Today, over 63% of all trips within the city center are made '
        'by bicycle — a figure that city planners in London, Paris, and Los Angeles '
        'can only dream of. The secret lies not just in building bike lanes, but in '
        'reshaping the entire urban environment around human-scale mobility.'
    )

    doc.add_paragraph(
        'Electric autonomous vehicles present another dimension of urban transformation. '
        'Companies like Waymo, Zoox, and Cruise are piloting self-driving taxi services '
        'in cities including San Francisco, Phoenix, and Austin. Proponents argue that '
        'autonomous vehicles could reduce urban parking requirements by up to 90%, '
        'freeing vast swaths of city land for parks, housing, and public spaces.'
    )

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['City', 'Cycling Share (%)', 'EV Adoption Rate (%)']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    city_data = [
        ('Amsterdam', '63', '28'),
        ('Copenhagen', '58', '35'),
        ('Oslo', '22', '82'),
        ('Paris', '15', '31'),
    ]
    for row_idx, (city, cycling, ev) in enumerate(city_data, 1):
        table.cell(row_idx, 0).text = city
        table.cell(row_idx, 1).text = cycling
        table.cell(row_idx, 2).text = ev

    doc.add_page_break()

    # ---- Page 4: Feature Article 3 ----
    doc.add_heading('Affordable Housing: Building for Everyone', level=1)

    doc.add_paragraph(
        'In cities from Hong Kong to San Francisco, the housing affordability crisis '
        'has reached a critical tipping point. Median home prices have outpaced wage '
        'growth for three consecutive decades, pushing working families and essential '
        'workers to the urban periphery or out of cities entirely.'
    )

    doc.add_paragraph(
        'Vienna\'s model of social housing, the Gemeindebauten, provides a counterpoint '
        'to this global trend. Today, approximately 60% of Vienna\'s 1.9 million '
        'residents live in subsidized housing — a legacy of progressive social policy '
        'dating back to the 1920s. The result is a city with one of the highest quality '
        'of life rankings in the world and a diverse, economically integrated populace.'
    )

    doc.add_heading('Innovative Housing Solutions Worldwide', level=2)

    solutions = [
        'Modular and prefabricated construction reducing costs by 20-30%',
        'Community Land Trusts permanently removing land from the speculative market',
        'Co-housing developments fostering intergenerational and mixed-income communities',
        'Adaptive reuse of commercial and industrial buildings for residential purposes',
        'Public-private partnerships leveraging institutional investment for affordable units',
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')

    doc.add_page_break()

    # ---- Page 5: Interview / Profile ----
    doc.add_heading('In Conversation: Dr. Maya Osei on Urban Resilience', level=1)

    doc.add_paragraph(
        'Dr. Maya Osei is a leading urban resilience researcher at the Global Cities '
        'Institute in Toronto. Her work focuses on helping cities prepare for and '
        'recover from climate-related shocks, from extreme heat events to catastrophic '
        'flooding. We sat down with her to discuss the future of urban planning.'
    )

    doc.add_heading('What does urban resilience mean to you?', level=3)

    doc.add_paragraph(
        '"Urban resilience is about more than just bouncing back from a disaster," '
        'Dr. Osei explains. "It\'s about building cities that can absorb shocks, adapt '
        'to changing conditions, and transform in ways that reduce future vulnerabilities. '
        'The cities that will thrive in the 21st century are those that treat resilience '
        'not as an emergency response, but as a core principle of everyday planning."'
    )

    doc.add_heading('Which cities are leading the way?', level=3)

    doc.add_paragraph(
        '"Rotterdam immediately comes to mind. After devastating floods in 1953, '
        'the Dutch completely transformed their approach to water management. Today, '
        'Rotterdam has water plazas that function as parks in dry weather but collect '
        'and store floodwater during heavy rain. It\'s a beautiful example of designing '
        'infrastructure to serve multiple purposes simultaneously."'
    )

    doc.add_page_break()

    # ---- Page 6: Data & Conclusions ----
    doc.add_heading('The Urban Data Revolution', level=1)

    doc.add_paragraph(
        'Smart city technologies are generating unprecedented volumes of data about '
        'how urban environments function. Sensor networks monitor air quality, '
        'pedestrian flows, energy consumption, and waste generation in real time. '
        'The challenge for city administrators is transforming this data torrent '
        'into actionable insights that improve daily life for residents.'
    )

    doc.add_paragraph(
        'Barcelona\'s Superblock initiative, launched in 2016, used granular traffic '
        'and noise data to redesign 503 city blocks, reducing car access and reclaiming '
        'streets for pedestrians and cyclists. Studies show that the project has reduced '
        'nitrogen dioxide concentrations by 24% in affected areas and contributed to '
        'a measurable reduction in noise pollution — with significant public health benefits.'
    )

    doc.add_heading('Looking Ahead', level=2)

    doc.add_paragraph(
        'The cities of tomorrow will not emerge from any single technology or planning '
        'philosophy. They will be built through countless small decisions — the choice '
        'to plant a tree rather than pour concrete, to zone for mixed uses rather than '
        'segregation, to invest in transit rather than highways. Each decision is a vote '
        'for the kind of urban future we want to inhabit.'
    )

    doc.add_paragraph(
        'What unites the best examples of urban innovation worldwide is a commitment '
        'to public space, equity, and long-term thinking. The cities profiled in this '
        'report demonstrate that it is possible to build urban environments that are '
        'simultaneously more productive, more sustainable, and more humane. The question '
        'is whether political will and public investment can keep pace with the urgency '
        'of the challenges ahead.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
