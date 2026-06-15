"""
Initial Setup: Magazine article with two-column layout, hyphenation disabled
Task ID: writer_pd_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# --- Magazine article content (long paragraphs with long words to show ragged edges) ---

TITLE = "The Transformative Power of Biotechnology in Contemporary Agriculture"

SUBTITLE = "How Cutting-Edge Scientific Breakthroughs Are Revolutionizing Global Food Production and Environmental Sustainability"

BODY_PARAGRAPHS = [
    # Section 1
    "The intersection of biotechnology and agricultural science represents one of the most consequential developments in modern civilization. Throughout the twentieth century, agricultural methodologies underwent unprecedented transformation, transitioning from predominantly manual labor-intensive practices to mechanized, scientifically sophisticated operations. Contemporary biotechnological advancements have further accelerated this metamorphosis, introducing revolutionary capabilities that fundamentally alter our understanding of crop development, pest management, and environmental stewardship.",

    "Genetic modification techniques, particularly the implementation of CRISPR-Cas9 gene-editing technology, have demonstrated extraordinary potential for enhancing crop characteristics. Researchers at prominent agricultural universities and biotechnology corporations have successfully developed drought-resistant wheat varieties, disease-tolerant rice cultivars, and nutritionally enhanced soybean strains that address critical food security challenges confronting developing nations.",

    "The environmental implications of these biotechnological interventions are multifaceted and occasionally controversial. Proponents emphasize the substantial reduction in chemical pesticide application, decreased water consumption requirements, and improved photosynthetic efficiency that characterize genetically optimized crop varieties. Environmentalists, conversely, express apprehension regarding potential ecological disruptions, biodiversity deterioration, and the unpredictable consequences of introducing genetically modified organisms into established ecosystems.",

    # Section 2
    "Dr. Margarethe Christopherson, a distinguished professor of agricultural biotechnology at the University of Wisconsin-Madison, has published extensively on the pharmacological applications of plant-derived biomolecules. Her groundbreaking research demonstrates that bioengineered photosynthetic organisms can simultaneously produce pharmaceutical compounds while maintaining their primary agricultural functionality, a phenomenon she describes as 'multifunctional bioarchitecture.'",

    "The commercialization of biotechnological agricultural products has precipitated substantial economic restructuring throughout the international agricultural marketplace. Multinational biotechnology conglomerates have established comprehensive intellectual property portfolios encompassing genetically modified seed varieties, proprietary cultivation methodologies, and specialized agrochemical formulations. This concentration of technological ownership has generated considerable controversy regarding accessibility, affordability, and the preservation of traditional agricultural knowledge systems.",

    "Developing nations, particularly those in sub-Saharan Africa and Southeast Asia, confront unique challenges in integrating biotechnological innovations into their agricultural frameworks. Infrastructure limitations, insufficient technical expertise, and regulatory uncertainties frequently impede the implementation of sophisticated biotechnological solutions. International development organizations, including the International Fund for Agricultural Development and the Consultative Group on International Agricultural Research, have established collaborative programs designed to facilitate technology transfer and capacity-building initiatives.",

    # Section 3
    "The microbiological dimensions of agricultural biotechnology represent another extraordinarily promising frontier. Soil microbiome analysis, utilizing advanced metagenomic sequencing technologies, has revealed previously uncharacterized symbiotic relationships between plant root systems and beneficial bacterial communities. Understanding and manipulating these microorganism-plant interactions offers unprecedented opportunities for enhancing nutrient absorption efficiency and natural disease suppression mechanisms.",

    "Precision agriculture, an interdisciplinary methodology integrating biotechnological innovations with sophisticated computational technologies, represents the contemporary convergence of agricultural science and digital infrastructure. Satellite-based monitoring systems, artificial intelligence-driven analytical frameworks, and Internet of Things sensor networks collectively enable extraordinarily precise management of agricultural operations, optimizing resource allocation while simultaneously minimizing environmental impact.",

    "The implementation of biopesticide technologies illustrates the practical application of biotechnological principles in sustainable pest management. Bacillus thuringiensis-derived proteins, mycoinsecticide formulations, and pheromone-based behavioral manipulation strategies provide environmentally responsible alternatives to conventional synthetic chemical pesticides. Agricultural entomologists have documented significant improvements in beneficial insect population maintenance and pollinator conservation following the adoption of these biotechnologically derived pest control methodologies.",

    # Section 4
    "Aquacultural biotechnology has emerged as a critically important subdiscipline, addressing the escalating global demand for protein-rich food sources. Genetically enhanced fish species, including growth-optimized Atlantic salmon and disease-resistant tilapia varieties, demonstrate the applicability of terrestrial biotechnological methodologies to aquatic food production systems. The controversial AquAdvantage salmon, representing the first genetically modified animal approved for human consumption, exemplifies both the extraordinary potential and the sociopolitical complexities associated with biotechnological innovation in food production.",

    "Phytoremediation research, leveraging the natural bioaccumulation capabilities of certain plant species, offers promising biotechnological solutions for environmental contamination challenges. Hyperaccumulator plants, genetically enhanced to improve their heavy metal absorption characteristics, can effectively extract toxic substances from contaminated agricultural soils, simultaneously rehabilitating degraded farmland and producing commercially valuable metal-enriched biomass.",

    "The philosophical and ethical considerations surrounding agricultural biotechnology necessitate comprehensive interdisciplinary examination. Questions regarding biodiversity preservation, intergenerational environmental responsibility, and the appropriate boundaries of technological intervention in natural biological processes demand thoughtful deliberation from scientists, policymakers, ethicists, and agricultural practitioners. The establishment of robust regulatory frameworks, incorporating transparent risk assessment methodologies and meaningful public participation mechanisms, remains essential for ensuring that biotechnological advancements in agriculture serve the collective interests of global society.",

    # Section 5
    "Nutritional enhancement through biofortification represents a particularly compelling application of agricultural biotechnology with direct humanitarian implications. The development of provitamin A-enriched Golden Rice, iron-fortified pearl millet, and zinc-enhanced wheat varieties demonstrates the extraordinary potential of biotechnological approaches to address micronutrient deficiency disorders affecting approximately two billion individuals worldwide. These biofortified crop varieties, developed through both conventional breeding accelerated by molecular marker-assisted selection and direct genetic modification techniques, offer cost-effective nutritional interventions for populations with limited dietary diversification opportunities.",

    "The convergence of nanotechnology and agricultural biotechnology has generated innovative approaches to nutrient delivery, pest detection, and environmental monitoring. Nanoencapsulated fertilizer formulations enable controlled-release nutrient distribution, substantially improving application efficiency while reducing environmental contamination associated with conventional fertilization practices. Biosensor nanotechnologies, incorporating genetically engineered biological recognition elements, facilitate real-time monitoring of pathogen presence, soil nutrient concentrations, and atmospheric conditions affecting crop development.",

    "Looking toward the future, the trajectory of agricultural biotechnology suggests increasingly sophisticated integration of biological engineering, computational intelligence, and environmental science. Synthetic biology approaches, enabling the design and construction of entirely novel biological systems, promise revolutionary capabilities for agricultural applications. From nitrogen-fixing cereals that eliminate dependence on manufactured fertilizers to perennial grain varieties that dramatically reduce soil erosion and cultivation energy requirements, the prospective contributions of biotechnological innovation to sustainable agriculture remain extraordinarily promising.",

    "The democratization of biotechnological capabilities through open-source research platforms, community laboratory initiatives, and decentralized innovation networks represents a transformative development in agricultural science. Citizen scientists, smallholder farmers, and independent researchers increasingly contribute to biotechnological knowledge generation, challenging the historical concentration of innovation capacity within institutional and corporate establishments. This participatory approach to agricultural biotechnology development offers promising pathways for ensuring that technological advancements address the diverse needs and circumstances of agricultural communities worldwide.",
]

HEADING_INDICES = {0: "Foundations of Agricultural Biotechnology",
                   3: "Research and Economic Perspectives",
                   6: "Microbiological and Digital Frontiers",
                   9: "Aquaculture and Environmental Applications",
                   12: "Nutritional Innovation and Future Horizons"}


def create_initial():
    doc = Document()

    # --- Page Setup: Two-column layout ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Set two columns via XML
    sectPr = section._sectPr
    cols_elem = sectPr.find(qn('w:cols'))
    if cols_elem is None:
        cols_elem = sectPr.makeelement(qn('w:cols'), {})
        sectPr.append(cols_elem)
    cols_elem.set(qn('w:num'), '2')
    cols_elem.set(qn('w:space'), '360')  # 0.25 inch gap between columns

    # --- Title ---
    title_para = doc.add_heading(TITLE, level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_para.paragraph_format.space_after = Pt(18)
    run = subtitle_para.add_run(SUBTITLE)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Author line ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_para.paragraph_format.space_after = Pt(24)
    run = author_para.add_run("By Dr. Eleanor Whitfield  |  March 2026  |  Agricultural Science Quarterly")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # --- Body paragraphs with section headings ---
    for i, para_text in enumerate(BODY_PARAGRAPHS):
        # Insert section heading if applicable
        if i in HEADING_INDICES:
            heading = doc.add_heading(HEADING_INDICES[i], level=2)
            # Ensure heading has NO hyphenation (explicitly)
            # Headings use "Heading 2" style, body uses "Normal"

        p = doc.add_paragraph(para_text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    # --- Ensure hyphenation is DISABLED for all paragraphs ---
    # We do NOT add any hyphenation XML elements, which means hyphenation is off by default

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
