"""
Reward Script: Automatic hyphenation settings for Magazine_Article.docx
Task ID: writer_pd_038
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Document-level autoHyphenation is enabled
  Component 2 (0.20): consecutiveHyphenLimit is set to 3
  Component 3 (0.30): Body text paragraphs have suppressAutoHyphens=0 (hyphenation enabled)
  Component 4 (0.20): Heading/Title paragraphs have suppressAutoHyphens=1 (hyphenation suppressed)
"""

import os
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_038'
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': WNS}


def verify_task(file_path):
    """
    Verify automatic hyphenation settings with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    settings_elem = doc.settings.element

    # Component 1: Document-level autoHyphenation is enabled (0.30 points)
    # In initial_env this element does not exist; in golden_env it is present.
    try:
        auto_hyph = settings_elem.find('w:autoHyphenation', NS)
        if auto_hyph is not None:
            # The element exists. If val attribute is absent, it means "true" by default.
            val = auto_hyph.attrib.get(f'{{{WNS}}}val', None)
            # No val attribute or val not "0"/"false" means enabled
            if val is None or val not in ('0', 'false'):
                print(f"PASS: Component 1 — autoHyphenation is enabled (val={val}) (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 1 — autoHyphenation element exists but is disabled (val={val})")
        else:
            print("FAIL: Component 1 — autoHyphenation element not found in document settings")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: consecutiveHyphenLimit is set to 3 (0.20 points)
    # In initial_env this element does not exist; in golden_env it is set to 3.
    try:
        consec = settings_elem.find('w:consecutiveHyphenLimit', NS)
        if consec is not None:
            val = consec.attrib.get(f'{{{WNS}}}val', None)
            if val == '3':
                print(f"PASS: Component 2 — consecutiveHyphenLimit is 3 (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — consecutiveHyphenLimit is {val}, expected 3")
        else:
            print("FAIL: Component 2 — consecutiveHyphenLimit element not found in document settings")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body text paragraphs have suppressAutoHyphens=0 (0.30 points)
    # In initial_env, paragraphs have no suppressAutoHyphens attribute.
    # In golden_env, Normal-styled body text paragraphs explicitly have suppressAutoHyphens="0".
    # We check that at least 50% of Normal-styled non-empty paragraphs have it set to "0".
    try:
        body_paras = []
        body_hyph_enabled = 0
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else 'None'
            # Only check Normal / body text style paragraphs with actual content
            if style_name in ('Normal', 'Body Text') and para.text.strip():
                body_paras.append(para)
                ppr = para._element.find('w:pPr', NS)
                if ppr is not None:
                    supp_el = ppr.find('w:suppressAutoHyphens', NS)
                    if supp_el is not None:
                        supp_val = supp_el.attrib.get(f'{{{WNS}}}val', '1')
                        if supp_val == '0':
                            body_hyph_enabled += 1

        if len(body_paras) == 0:
            print("FAIL: Component 3 — no body text paragraphs found")
        else:
            ratio = body_hyph_enabled / len(body_paras)
            if ratio >= 0.5:
                print(f"PASS: Component 3 — {body_hyph_enabled}/{len(body_paras)} body paragraphs have hyphenation enabled (ratio={ratio:.2f}) (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 3 — only {body_hyph_enabled}/{len(body_paras)} body paragraphs have suppressAutoHyphens=0 (ratio={ratio:.2f})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Heading/Title paragraphs have suppressAutoHyphens=1 (0.20 points)
    # In initial_env, headings have no suppressAutoHyphens attribute.
    # In golden_env, Heading and Title paragraphs have suppressAutoHyphens="1".
    # We check that at least 50% of heading/title paragraphs have it set to "1".
    try:
        heading_paras = []
        heading_hyph_suppressed = 0
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else 'None'
            if ('Heading' in style_name or style_name == 'Title') and para.text.strip():
                heading_paras.append(para)
                ppr = para._element.find('w:pPr', NS)
                if ppr is not None:
                    supp_el = ppr.find('w:suppressAutoHyphens', NS)
                    if supp_el is not None:
                        supp_val = supp_el.attrib.get(f'{{{WNS}}}val', '1')
                        if supp_val == '1':
                            heading_hyph_suppressed += 1

        if len(heading_paras) == 0:
            print("FAIL: Component 4 — no heading/title paragraphs found")
        else:
            ratio = heading_hyph_suppressed / len(heading_paras)
            if ratio >= 0.5:
                print(f"PASS: Component 4 — {heading_hyph_suppressed}/{len(heading_paras)} heading paragraphs have hyphenation suppressed (ratio={ratio:.2f}) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — only {heading_hyph_suppressed}/{len(heading_paras)} heading paragraphs have suppressAutoHyphens=1 (ratio={ratio:.2f})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
