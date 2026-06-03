"""
Initial Setup: Complex document with three page style sections (Cover, TOC, Body)
Task ID: writer_rd_095
Domain: libreoffice_writer

Creates a 30-page Writer document with uniform Default Page Style:
- Page 1: Cover page
- Pages 2-3: Table of contents
- Pages 4-30: Main body content (chapters)
- All pages: Arabic numbering 1-30, header 'Technical Manual', uniform 2.54 cm margins
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_095'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_field(paragraph):
    """Add a PAGE field code to a paragraph."""
    r1 = paragraph.add_run()
    fldChar1 = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fldChar1)

    r2 = paragraph.add_run()
    instrText = r2._element.makeelement(qn('w:instrText'), {})
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    r2._element.append(instrText)

    r3 = paragraph.add_run()
    fldChar2 = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fldChar2)


def add_page_break_in_paragraph(para):
    """Add a page break at the end of a paragraph."""
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # -- Set uniform margins: 2.54 cm (1 inch) all around --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Set header 'Technical Manual' on all sections --
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = 'Technical Manual'
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = hp.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # -- Set footer with Arabic page number --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number_field(fp)

    # ========== PAGE 1: COVER PAGE ==========
    # Add some blank space before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    title = doc.add_heading('Nextera Solutions', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('Technical Infrastructure Manual', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ver = doc.add_paragraph()
    ver.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = ver.add_run('Version 4.2 — March 2025')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    conf = doc.add_paragraph()
    conf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = conf.add_run('CONFIDENTIAL')
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    r2.bold = True

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r3 = dept.add_run('Prepared by the Infrastructure Engineering Division')
    r3.font.size = Pt(11)

    # Page break after cover
    add_page_break_in_paragraph(dept)

    # ========== PAGES 2-3: TABLE OF CONTENTS ==========
    toc_title = doc.add_heading('Table of Contents', level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    toc_entries = [
        ('1. Network Architecture Overview', '4'),
        ('   1.1 Core Network Design', '5'),
        ('   1.2 VLAN Segmentation Strategy', '6'),
        ('   1.3 Firewall Rules and Policies', '7'),
        ('2. Server Infrastructure', '8'),
        ('   2.1 Hypervisor Cluster Configuration', '9'),
        ('   2.2 Storage Area Network Layout', '10'),
        ('   2.3 Backup and Disaster Recovery', '11'),
        ('3. Cloud Integration', '12'),
        ('   3.1 AWS Multi-Region Setup', '13'),
        ('   3.2 Azure AD Federation', '14'),
        ('   3.3 GCP BigQuery Data Pipeline', '15'),
        ('4. Security Framework', '16'),
        ('   4.1 Zero Trust Architecture', '17'),
        ('   4.2 Identity and Access Management', '18'),
        ('   4.3 Encryption Standards', '19'),
        ('5. Monitoring and Alerting', '20'),
        ('   5.1 Prometheus and Grafana Stack', '21'),
        ('   5.2 Log Aggregation with ELK', '22'),
        ('   5.3 Incident Response Procedures', '23'),
        ('6. Deployment Pipelines', '24'),
        ('   6.1 CI/CD with Jenkins and ArgoCD', '25'),
        ('   6.2 Container Orchestration (Kubernetes)', '26'),
        ('   6.3 Blue-Green Deployment Strategy', '27'),
        ('7. Compliance and Auditing', '28'),
        ('   7.1 SOC 2 Type II Controls', '29'),
        ('   7.2 GDPR Data Processing Records', '30'),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        r = p.add_run(f'{entry} ')
        r.font.size = Pt(11)
        dots = '.' * (60 - len(entry))
        rd = p.add_run(dots)
        rd.font.size = Pt(11)
        rd.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        rp = p.add_run(f' {page}')
        rp.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    # Page break after TOC (end of page 3)
    last_toc = doc.paragraphs[-1]
    add_page_break_in_paragraph(last_toc)

    # ========== PAGES 4-30: MAIN BODY CONTENT ==========
    chapters = [
        {
            'title': 'Chapter 1: Network Architecture Overview',
            'sections': [
                ('1.1 Core Network Design',
                 'The core network at Nextera Solutions is built on a spine-leaf architecture using Cisco Nexus 9000 series switches. '
                 'The spine layer consists of four Nexus 9504 chassis providing 400 Gbps aggregate bandwidth between leaf switches. '
                 'Each leaf switch connects to a maximum of 48 server ports at 25 Gbps, with dual uplinks to every spine for redundancy. '
                 'BGP is used as the underlay routing protocol with EVPN-VXLAN overlay for multi-tenancy.\n\n'
                 'Traffic engineering policies ensure that latency-sensitive applications like the real-time trading platform '
                 'receive priority queuing via DSCP marking at ingress. The mean latency across the fabric is 3.2 microseconds '
                 'under normal load, scaling to 8.7 microseconds at 80% utilization.'),
                ('1.2 VLAN Segmentation Strategy',
                 'Production workloads are segmented into 12 VLANs across three security zones. Zone A (VLANs 100-103) hosts '
                 'customer-facing services including the API gateway cluster and web front-ends. Zone B (VLANs 200-205) contains '
                 'internal application services, message queues, and caching layers. Zone C (VLANs 300-302) is reserved for '
                 'database clusters and persistent storage.\n\n'
                 'Inter-zone traffic passes through Palo Alto PA-5250 firewalls with application-level inspection enabled. '
                 'East-west traffic within Zone B is monitored by Gigamon visibility nodes feeding into the Darktrace AI engine '
                 'for anomaly detection.'),
                ('1.3 Firewall Rules and Policies',
                 'The firewall ruleset follows the principle of least privilege with 847 active rules as of Q1 2025. '
                 'Rule reviews are conducted quarterly by the Security Operations team led by Diana Kowalski. '
                 'Deprecated rules older than 180 days without traffic hits are automatically flagged for removal.\n\n'
                 'Key policies include: TLS 1.3 enforcement on all external connections, SSH key-only authentication '
                 'for management interfaces, and IP reputation filtering using Talos Intelligence feeds updated every 15 minutes.'),
            ]
        },
        {
            'title': 'Chapter 2: Server Infrastructure',
            'sections': [
                ('2.1 Hypervisor Cluster Configuration',
                 'The primary data center operates 64 VMware ESXi 8.0 hosts organized into four clusters. '
                 'Cluster Alpha handles production workloads with 16 hosts, each equipped with dual AMD EPYC 9654 processors '
                 '(96 cores per socket), 2 TB DDR5 RAM, and 8x NVMe drives in RAID-10. DRS is configured with a migration '
                 'threshold of 3 to balance CPU and memory utilization.\n\n'
                 'Resource pools enforce hard limits: the trading platform receives a guaranteed 256 vCPUs and 1 TB RAM, '
                 'while batch processing workloads are capped at 50% of remaining cluster capacity during business hours.'),
                ('2.2 Storage Area Network Layout',
                 'The SAN fabric uses dual Brocade G630 switches with ISL trunking for 128 Gbps inter-switch bandwidth. '
                 'Pure Storage FlashArray//X70 provides 500 TB of usable NVMe storage with consistent sub-millisecond latency. '
                 'Tiered storage policies automatically move infrequently accessed data to NetApp FAS9500 systems with 2 PB '
                 'of hybrid SSD/HDD capacity.\n\n'
                 'Snapshot schedules run hourly for production volumes, with daily snapshots retained for 30 days and weekly '
                 'snapshots retained for 12 months. Replication to the disaster recovery site in Frankfurt operates '
                 'asynchronously with an RPO of 15 minutes.'),
                ('2.3 Backup and Disaster Recovery',
                 'Veeam Backup & Replication v12 manages all backup operations. Full backups run weekly on Sundays with '
                 'incremental backups every 6 hours. The backup repository consists of 12 Dell PowerScale nodes providing '
                 '800 TB of deduplication-enabled storage, achieving a 14:1 deduplication ratio on average.\n\n'
                 'Disaster recovery testing occurs quarterly under the supervision of Operations Manager Raj Patel. '
                 'The last DR test in February 2025 achieved an RTO of 47 minutes for Tier 1 services, well within the '
                 '60-minute SLA. Tier 2 services recovered in 2 hours 15 minutes against a 4-hour target.'),
            ]
        },
        {
            'title': 'Chapter 3: Cloud Integration',
            'sections': [
                ('3.1 AWS Multi-Region Setup',
                 'Nextera maintains active-active deployments across us-east-1 (Virginia) and eu-west-1 (Ireland). '
                 'Route 53 health checks with 10-second intervals handle automatic failover. The Virginia region hosts '
                 '47 EC2 instances across three availability zones, while Ireland runs 32 instances for European traffic.\n\n'
                 'Cost optimization through Reserved Instances and Savings Plans covers 72% of compute spend. '
                 'Spot instances handle batch ML training jobs, reducing compute costs by 65% compared to on-demand pricing. '
                 'Monthly AWS spend averages $187,000 with a 12-month trend showing 8% cost reduction despite 23% traffic growth.'),
                ('3.2 Azure AD Federation',
                 'Single sign-on is provided through Azure AD Premium P2 with SAML 2.0 federation to 28 internal applications. '
                 'Conditional access policies enforce MFA for all users, with hardware security keys required for privileged accounts. '
                 'The identity team manages 2,847 user accounts and 156 service principals.\n\n'
                 'Privileged Identity Management (PIM) governs just-in-time access to 42 administrative roles. '
                 'Activation requires business justification and approval from the role owner. Average activation duration '
                 'is 4 hours, with a maximum of 8 hours for most roles.'),
                ('3.3 GCP BigQuery Data Pipeline',
                 'The analytics platform ingests 2.3 TB of data daily into BigQuery through Dataflow streaming pipelines. '
                 'Source systems include the production PostgreSQL cluster (via Debezium CDC), Salesforce (via Fivetran), '
                 'and IoT sensor data from 14,000 edge devices (via Pub/Sub).\n\n'
                 'Data transformation uses dbt with 340 models organized in a medallion architecture. Bronze tables retain '
                 'raw data for 90 days, silver tables apply business logic and deduplication, and gold tables serve analytics '
                 'dashboards accessed by 180 active users through Looker.'),
            ]
        },
        {
            'title': 'Chapter 4: Security Framework',
            'sections': [
                ('4.1 Zero Trust Architecture',
                 'The zero trust implementation follows NIST SP 800-207 guidelines with continuous verification at every access point. '
                 'Zscaler Private Access replaces traditional VPN for remote workers, providing application-level micro-tunnels '
                 'authenticated by device posture and user identity. 94% of the workforce now accesses internal resources '
                 'through ZPA, with legacy VPN scheduled for decommission in Q3 2025.\n\n'
                 'Network micro-segmentation using VMware NSX-T enforces 2,340 firewall rules between workload segments. '
                 'East-west traffic inspection catches an average of 127 policy violations per month, primarily from '
                 'misconfigured development containers attempting unauthorized database connections.'),
                ('4.2 Identity and Access Management',
                 'CyberArk Privileged Access Security manages 892 privileged accounts across infrastructure components. '
                 'Password rotation occurs every 24 hours for service accounts and every 90 days for human administrators. '
                 'Session recording captures all privileged activities with 6-month retention for compliance purposes.\n\n'
                 'Access reviews run quarterly through SailPoint IdentityNow. The last review cycle in January 2025 '
                 'processed 12,400 entitlements, revoking 340 (2.7%) due to role changes or terminations. '
                 'The average review completion time was 8 business days across 47 application owners.'),
                ('4.3 Encryption Standards',
                 'All data at rest is encrypted with AES-256. TLS 1.3 is mandatory for data in transit, with legacy '
                 'TLS 1.2 connections blocked since January 2025. Certificate management uses HashiCorp Vault with '
                 'automatic rotation for internal PKI certificates every 90 days.\n\n'
                 'Key management follows a hierarchical model: HSM-protected master keys in Thales Luna 7 appliances, '
                 'data encryption keys derived per-service, and ephemeral session keys for real-time communications. '
                 'The cryptographic inventory tracks 4,200 active keys across all systems.'),
            ]
        },
        {
            'title': 'Chapter 5: Monitoring and Alerting',
            'sections': [
                ('5.1 Prometheus and Grafana Stack',
                 'The monitoring infrastructure collects 2.8 million metrics per minute across 847 targets. '
                 'Prometheus runs in a highly available pair with Thanos for long-term storage and global querying. '
                 'Retention is 15 days locally and 13 months in object storage (MinIO cluster with 120 TB capacity).\n\n'
                 'Grafana serves 142 dashboards organized by team and service. The SRE team maintains 67 alerting rules '
                 'with a current false positive rate of 3.2%. PagerDuty integration routes critical alerts to the on-call '
                 'engineer with a median acknowledgment time of 4 minutes during business hours.'),
                ('5.2 Log Aggregation with ELK',
                 'The Elasticsearch cluster comprises 18 data nodes with 4 TB NVMe storage each, processing 450 GB of '
                 'logs daily. Index lifecycle management retains hot data for 7 days, warm for 30 days, and cold in '
                 'S3 for 12 months. Logstash pipelines parse 28 distinct log formats with Grok patterns maintained '
                 'by the Platform Engineering team.\n\n'
                 'Kibana dashboards provide real-time visibility into application performance, security events, and '
                 'infrastructure health. Custom Machine Learning jobs detect unusual log patterns, generating an average '
                 'of 15 anomaly alerts per week, 60% of which lead to actionable investigations.'),
                ('5.3 Incident Response Procedures',
                 'Incidents are classified into four severity levels. SEV-1 incidents require the Incident Commander '
                 '(rotating weekly among 8 senior engineers) to assemble a war room within 15 minutes. Communication '
                 'flows through a dedicated Slack channel with automated status page updates via Statuspage.io.\n\n'
                 'Post-incident reviews are mandatory for SEV-1 and SEV-2 events. The blameless retrospective format '
                 'includes timeline reconstruction, root cause analysis, and action items tracked in Jira. '
                 'In Q1 2025, the team resolved 3 SEV-1 and 11 SEV-2 incidents, with MTTR of 38 minutes and '
                 '2.1 hours respectively.'),
            ]
        },
        {
            'title': 'Chapter 6: Deployment Pipelines',
            'sections': [
                ('6.1 CI/CD with Jenkins and ArgoCD',
                 'The build pipeline processes an average of 340 builds per day across 86 microservices. '
                 'Jenkins orchestrates compilation, unit testing, and artifact packaging. SonarQube quality gates '
                 'enforce a minimum 80% code coverage, zero critical vulnerabilities, and less than 3% code duplication.\n\n'
                 'ArgoCD manages GitOps-based deployment to Kubernetes clusters. Application definitions in Helm charts '
                 'are version-controlled in a dedicated deployment repository. Sync waves ensure database migrations '
                 'complete before application pods roll out, preventing schema mismatch errors.'),
                ('6.2 Container Orchestration (Kubernetes)',
                 'Three Kubernetes clusters serve distinct purposes: production (72 nodes), staging (24 nodes), and '
                 'development (12 nodes). All clusters run on Rancher-managed RKE2 with CIS-hardened node configurations. '
                 'Pod security standards enforce restricted profiles for all namespaces except system components.\n\n'
                 'Resource quotas prevent namespace sprawl: each team receives CPU and memory limits proportional to '
                 'their service tier. Horizontal Pod Autoscaler manages 68% of production deployments, with KEDA '
                 'handling event-driven scaling for message queue consumers.'),
                ('6.3 Blue-Green Deployment Strategy',
                 'Critical services use blue-green deployment through Istio service mesh traffic splitting. '
                 'The deployment process: green environment receives new version, automated smoke tests validate '
                 'health endpoints and critical user journeys, traffic shifts 10% then 50% then 100% with '
                 '5-minute observation windows between steps.\n\n'
                 'Rollback triggers automatically if error rate exceeds 0.1% or p99 latency increases by more than '
                 '20% during canary observation. In 2024, 12 out of 890 deployments triggered automatic rollback, '
                 'preventing customer-facing impact in all cases.'),
            ]
        },
        {
            'title': 'Chapter 7: Compliance and Auditing',
            'sections': [
                ('7.1 SOC 2 Type II Controls',
                 'Nextera Solutions maintains SOC 2 Type II certification audited annually by Deloitte. '
                 'The control framework covers 94 controls across security, availability, and confidentiality trust '
                 'service criteria. Continuous control monitoring through Drata automates evidence collection for '
                 '78% of controls, reducing audit preparation from 6 weeks to 2 weeks.\n\n'
                 'Key controls include: CC6.1 (logical access), CC6.6 (encryption), CC7.2 (system monitoring), '
                 'and CC8.1 (change management). The most recent audit in November 2024 resulted in zero exceptions, '
                 'maintaining the clean report streak for the fourth consecutive year.'),
                ('7.2 GDPR Data Processing Records',
                 'The data protection team led by DPO Elena Vasquez maintains processing records for 34 data processing '
                 'activities involving EU personal data. Data protection impact assessments have been completed for '
                 '12 high-risk processing operations including the customer analytics platform and employee monitoring system.\n\n'
                 'Data subject access requests are processed through a self-service portal with automated identity '
                 'verification. The average DSAR fulfillment time is 12 business days against the 30-day regulatory '
                 'deadline. In 2024, the team processed 247 DSARs with 100% on-time completion.'),
            ]
        },
    ]

    page_count = 4  # Starting at page 4
    for ch_idx, chapter in enumerate(chapters):
        # Chapter title
        h = doc.add_heading(chapter['title'], level=1)

        for sec_idx, (sec_title, sec_content) in enumerate(chapter['sections']):
            doc.add_heading(sec_title, level=2)
            paragraphs = sec_content.split('\n\n')
            for para_text in paragraphs:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(8)

            page_count += 1
            # Add page break between sections (except the very last section)
            if not (ch_idx == len(chapters) - 1 and sec_idx == len(chapter['sections']) - 1):
                doc.add_page_break()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
