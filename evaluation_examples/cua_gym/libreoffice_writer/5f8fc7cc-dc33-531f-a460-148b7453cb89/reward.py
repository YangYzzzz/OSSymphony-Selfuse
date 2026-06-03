"""
Reward Script: Complex page styles (Cover / TOC / Body) for a Writer document
Task ID: writer_rd_095
Domain: libreoffice_writer
Scoring:
  C1 (0.15) — Document has 3 sections (Cover, TOC, Body)
  C2 (0.25) — Cover section: no header/footer text, ~4cm top margin, light-blue shading (#E6F0FF)
  C3 (0.20) — TOC section: 'Contents' centered italic header, Roman numeral page numbers starting at 1
  C4 (0.20) — Body section: mirrored margins (~3cm inner, ~2cm outer), Arabic page numbers restarted at 1
  C5 (0.20) — Body section: even header='Technical Manual', odd header has chapter text
"""

import os
import zipfile
import xml.etree.ElementTree as ET

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_095'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def persist_app_state(domain):
    """Try to save any unsaved document via Ctrl+S."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """Verify task completion with progressive scoring. Returns float 0.0-1.0."""
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)

    # ── Component 1: Document has 3 sections (0.15 pts) ──
    try:
        if num_sections >= 3:
            print(f"PASS: C1 — Document has {num_sections} sections (>= 3) (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: C1 — Expected >= 3 sections, found {num_sections}")
    except Exception as e:
        print(f"ERROR: C1 — {e}")

    if num_sections < 3:
        # Cannot verify further without 3 sections
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    sec_cover = doc.sections[0]
    sec_toc = doc.sections[1]
    sec_body = doc.sections[2]

    # ── Component 2: Cover section properties (0.25 pts) ──
    # Sub-checks: no header text (0.05), no footer text (0.05), ~4cm top margin (0.05), E6F0FF shading (0.10)
    try:
        c2_score = 0.0

        # 2a: No header text in cover
        cover_hdr_text = ""
        if not sec_cover.header.is_linked_to_previous:
            cover_hdr_text = " ".join(p.text.strip() for p in sec_cover.header.paragraphs).strip()
        if cover_hdr_text == "":
            print("PASS: C2a — Cover has no header text (0.05 pts)")
            c2_score += 0.05
        else:
            print(f"FAIL: C2a — Cover header has text: '{cover_hdr_text}'")

        # 2b: No footer text in cover (no page numbers, no text)
        cover_ftr_text = ""
        cover_ftr_has_field = False
        if not sec_cover.footer.is_linked_to_previous:
            cover_ftr_text = " ".join(p.text.strip() for p in sec_cover.footer.paragraphs).strip()
            for p in sec_cover.footer.paragraphs:
                instr = p._element.findall('.//w:instrText', NS)
                if instr:
                    cover_ftr_has_field = True
        if cover_ftr_text == "" and not cover_ftr_has_field:
            print("PASS: C2b — Cover has no footer text/page numbers (0.05 pts)")
            c2_score += 0.05
        else:
            print(f"FAIL: C2b — Cover footer has text='{cover_ftr_text}', field={cover_ftr_has_field}")

        # 2c: Top margin ~4cm (4cm = 1440000 EMU, tolerance 10%)
        top_margin_cm = sec_cover.top_margin / 360000.0
        if 3.5 <= top_margin_cm <= 4.5:
            print(f"PASS: C2c — Cover top margin {top_margin_cm:.2f} cm (~4cm) (0.05 pts)")
            c2_score += 0.05
        else:
            print(f"FAIL: C2c — Cover top margin {top_margin_cm:.2f} cm, expected ~4cm")

        # 2d: Paragraphs in cover section have shading #E6F0FF
        # We check via XML: paragraphs before first section break should have shading
        try:
            zf = zipfile.ZipFile(file_path)
            doc_xml = ET.fromstring(zf.read('word/document.xml'))
            body = doc_xml.find('w:body', NS)
            cover_shaded = 0
            cover_total = 0
            for elem in body:
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                if tag == 'p':
                    cover_total += 1
                    pPr = elem.find('w:pPr', NS)
                    if pPr is not None:
                        sect = pPr.find('w:sectPr', NS)
                        if sect is not None:
                            break  # reached end of first section
                        shd = pPr.find('w:shd', NS)
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill.upper() == 'E6F0FF':
                                cover_shaded += 1
            zf.close()
            if cover_total > 0 and cover_shaded >= cover_total * 0.8:
                print(f"PASS: C2d — Cover has E6F0FF shading ({cover_shaded}/{cover_total} paras) (0.10 pts)")
                c2_score += 0.10
            else:
                print(f"FAIL: C2d — Cover shading: {cover_shaded}/{cover_total} paras with E6F0FF")
        except Exception as e:
            print(f"ERROR: C2d — {e}")

        total_score += c2_score
        print(f"  C2 subtotal: {c2_score}/0.25")
    except Exception as e:
        print(f"ERROR: C2 — {e}")

    # ── Component 3: TOC section properties (0.20 pts) ──
    # Sub-checks: header='Contents' centered italic (0.10), Roman numeral page numbers starting at 1 (0.10)
    try:
        c3_score = 0.0

        # 3a: Header with 'Contents' centered in italic
        toc_hdr_text = ""
        toc_hdr_italic = False
        toc_hdr_center = False
        if not sec_toc.header.is_linked_to_previous:
            for p in sec_toc.header.paragraphs:
                if p.text.strip():
                    toc_hdr_text = p.text.strip()
                    # Check alignment
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                        toc_hdr_center = True
                    # Check italic on runs
                    for run in p.runs:
                        if run.text.strip() and run.font.italic:
                            toc_hdr_italic = True

        if 'contents' in toc_hdr_text.lower() and toc_hdr_center and toc_hdr_italic:
            print(f"PASS: C3a — TOC header='Contents', centered, italic (0.10 pts)")
            c3_score += 0.10
        else:
            print(f"FAIL: C3a — TOC header text='{toc_hdr_text}', center={toc_hdr_center}, italic={toc_hdr_italic}")

        # 3b: Roman numeral page numbers starting at 1
        toc_sectPr = sec_toc._sectPr
        pgNumType = toc_sectPr.find(qn('w:pgNumType'))
        has_roman = False
        starts_at_1 = False
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            if fmt and 'roman' in fmt.lower():
                has_roman = True
            if start == '1':
                starts_at_1 = True

        # Also check footer has PAGE field
        toc_footer_has_page = False
        if not sec_toc.footer.is_linked_to_previous:
            for p in sec_toc.footer.paragraphs:
                for it in p._element.findall('.//w:instrText', NS):
                    if it.text and 'PAGE' in it.text.upper():
                        toc_footer_has_page = True

        if has_roman and starts_at_1 and toc_footer_has_page:
            print(f"PASS: C3b — TOC has Roman page numbers starting at 1 (0.10 pts)")
            c3_score += 0.10
        else:
            print(f"FAIL: C3b — Roman={has_roman}, start=1={starts_at_1}, footer_page={toc_footer_has_page}")

        total_score += c3_score
        print(f"  C3 subtotal: {c3_score}/0.20")
    except Exception as e:
        print(f"ERROR: C3 — {e}")

    # ── Component 4: Body section margins and page numbers (0.20 pts) ──
    # Sub-checks: mirrored margins ~3cm inner / ~2cm outer (0.10), Arabic numbers restarted at 1 (0.10)
    try:
        c4_score = 0.0

        # 4a: Mirrored margins: inner ~3cm, outer ~2cm
        left_cm = sec_body.left_margin / 360000.0
        right_cm = sec_body.right_margin / 360000.0
        # Inner margin (left) should be ~3cm, outer (right) ~2cm
        margins_ok = (2.5 <= left_cm <= 3.5) and (1.5 <= right_cm <= 2.5)
        # Also check mirrorMargins is set in settings.xml OR gutter margin
        try:
            zf2 = zipfile.ZipFile(file_path)
            settings_xml = ET.fromstring(zf2.read('word/settings.xml'))
            mirror_elem = settings_xml.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}mirrorMargins')
            has_mirror = mirror_elem is not None
            zf2.close()
        except:
            has_mirror = False

        if margins_ok:
            if has_mirror:
                print(f"PASS: C4a — Body margins: left={left_cm:.2f}cm, right={right_cm:.2f}cm, mirrorMargins=True (0.10 pts)")
                c4_score += 0.10
            else:
                # Accept asymmetric margins even without explicit mirrorMargins flag (partial)
                print(f"PASS: C4a — Body margins: left={left_cm:.2f}cm, right={right_cm:.2f}cm (asymmetric, no mirror flag) (0.07 pts)")
                c4_score += 0.07
        else:
            print(f"FAIL: C4a — Body margins: left={left_cm:.2f}cm, right={right_cm:.2f}cm, expected ~3cm/~2cm")

        # 4b: Arabic page numbers restarted at 1
        body_sectPr = sec_body._sectPr
        pgNumType = body_sectPr.find(qn('w:pgNumType'))
        body_arabic = False
        body_start_1 = False
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            if fmt is None or fmt == 'decimal':
                body_arabic = True
            if start == '1':
                body_start_1 = True

        body_footer_has_page = False
        if not sec_body.footer.is_linked_to_previous:
            for p in sec_body.footer.paragraphs:
                for it in p._element.findall('.//w:instrText', NS):
                    if it.text and 'PAGE' in it.text.upper():
                        body_footer_has_page = True

        if body_arabic and body_start_1 and body_footer_has_page:
            print(f"PASS: C4b — Body has Arabic page numbers starting at 1 (0.10 pts)")
            c4_score += 0.10
        else:
            print(f"FAIL: C4b — Arabic={body_arabic}, start=1={body_start_1}, footer_page={body_footer_has_page}")

        total_score += c4_score
        print(f"  C4 subtotal: {c4_score}/0.20")
    except Exception as e:
        print(f"ERROR: C4 — {e}")

    # ── Component 5: Body section headers (even/odd) (0.20 pts) ──
    # Sub-checks: even header='Technical Manual' (0.10), odd header has chapter text (0.10)
    try:
        c5_score = 0.0

        # 5a: Even page header has 'Technical Manual'
        even_hdr_text = ""
        if not sec_body.even_page_header.is_linked_to_previous:
            for p in sec_body.even_page_header.paragraphs:
                if p.text.strip():
                    even_hdr_text = p.text.strip()

        if 'technical manual' in even_hdr_text.lower():
            print(f"PASS: C5a — Body even header='{even_hdr_text}' (0.10 pts)")
            c5_score += 0.10
        else:
            print(f"FAIL: C5a — Body even header='{even_hdr_text}', expected 'Technical Manual'")

        # 5b: Odd (default) header has chapter-related text
        odd_hdr_text = ""
        if not sec_body.header.is_linked_to_previous:
            for p in sec_body.header.paragraphs:
                if p.text.strip():
                    odd_hdr_text = p.text.strip()

        # The odd header should contain chapter-related text (not just 'Technical Manual')
        # and should differ from the even header
        if odd_hdr_text and odd_hdr_text.lower() != even_hdr_text.lower():
            print(f"PASS: C5b — Body odd header='{odd_hdr_text}' (differs from even) (0.10 pts)")
            c5_score += 0.10
        else:
            print(f"FAIL: C5b — Body odd header='{odd_hdr_text}', should differ from even header")

        # Also verify evenAndOddHeaders is enabled in settings
        try:
            zf3 = zipfile.ZipFile(file_path)
            settings_xml = ET.fromstring(zf3.read('word/settings.xml'))
            eo = settings_xml.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}evenAndOddHeaders')
            zf3.close()
            if eo is not None:
                print(f"  INFO: evenAndOddHeaders is set in settings.xml")
            else:
                print(f"  WARN: evenAndOddHeaders not found in settings.xml (headers may not alternate)")
        except:
            pass

        total_score += c5_score
        print(f"  C5 subtotal: {c5_score}/0.20")
    except Exception as e:
        print(f"ERROR: C5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
