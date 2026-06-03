"""
Reward Script: Apply italic formatting to all bracketed translation notes
Task ID: writer_frd_031
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Proportion of 14 bracketed segments that are italic
  Component 2 (0.3): Brackets themselves are included in italic runs
  Component 3 (0.2): Non-bracketed text is not italic (no over-application)
"""

import os
import re
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_031'
EXPECTED_BRACKET_COUNT = 14


def persist_app_state(domain: str):
    """Save any unsaved changes in the GUI application."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    bracket_pattern = re.compile(r'\[.+?\]')

    # ---------------------------------------------------------------
    # Component 1: Proportion of bracketed segments that are italic
    # (0.5 points — progressive based on how many of the 14 are done)
    # ---------------------------------------------------------------
    try:
        italic_bracket_count = 0
        total_bracket_count = 0

        for para in doc.paragraphs:
            full_text = para.text
            matches = list(bracket_pattern.finditer(full_text))
            if not matches:
                continue

            for match in matches:
                total_bracket_count += 1
                bracket_text = match.group()

                # Check if this bracketed segment is in a run with italic=True
                # We need to find the run(s) that contain this bracket text
                found_italic = False
                for run in para.runs:
                    if bracket_text in run.text or (
                        '[' in run.text and run.text.strip().startswith('[')
                    ):
                        # This run contains our bracketed text
                        if run.font.italic is True:
                            found_italic = True
                            break

                if found_italic:
                    italic_bracket_count += 1

        if total_bracket_count == 0:
            print("FAIL: Component 1 — No bracketed segments found in document")
        else:
            ratio = italic_bracket_count / max(total_bracket_count, EXPECTED_BRACKET_COUNT)
            comp1_score = 0.5 * ratio
            total_score += comp1_score
            print(f"PASS: Component 1 — {italic_bracket_count}/{total_bracket_count} bracketed segments are italic ({comp1_score:.2f} pts)")
            if italic_bracket_count < total_bracket_count:
                print(f"  NOTE: {total_bracket_count - italic_bracket_count} bracketed segments still not italic")

    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: Brackets are included in the italic run
    # (0.3 points — verifies that [ and ] themselves are italic, not just inner text)
    # ---------------------------------------------------------------
    try:
        brackets_in_italic_run = 0
        total_brackets_checked = 0

        for para in doc.paragraphs:
            full_text = para.text
            matches = list(bracket_pattern.finditer(full_text))
            if not matches:
                continue

            for match in matches:
                total_brackets_checked += 1
                bracket_text = match.group()

                # Find run containing this exact bracketed text
                bracket_properly_italic = False
                for run in para.runs:
                    # The run should contain the full bracket text including [ and ]
                    if bracket_text in run.text and run.font.italic is True:
                        bracket_properly_italic = True
                        break

                if bracket_properly_italic:
                    brackets_in_italic_run += 1

        if total_brackets_checked > 0:
            ratio2 = brackets_in_italic_run / max(total_brackets_checked, EXPECTED_BRACKET_COUNT)
            comp2_score = 0.3 * ratio2
            total_score += comp2_score
            print(f"PASS: Component 2 — {brackets_in_italic_run}/{total_brackets_checked} brackets included in italic runs ({comp2_score:.2f} pts)")
        else:
            print("FAIL: Component 2 — No bracketed segments to check")

    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Non-bracketed text is NOT italic (no over-application)
    # (0.2 points — ensures formatting was targeted, not blanket)
    # ---------------------------------------------------------------
    try:
        non_bracket_runs_total = 0
        non_bracket_runs_italic = 0

        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text.strip()
                if not text:
                    continue
                # Skip runs that contain bracketed text
                if '[' in run.text or ']' in run.text:
                    continue
                non_bracket_runs_total += 1
                if run.font.italic is True:
                    non_bracket_runs_italic += 1

        if non_bracket_runs_total == 0:
            print("FAIL: Component 3 — No non-bracketed runs to check")
        elif italic_bracket_count == 0:
            # Gate: only award this component if at least some brackets were italicized
            # Otherwise this is a precondition (no italic anywhere), not a task change
            print(f"FAIL: Component 3 — No brackets are italic yet, so non-italic non-brackets is a precondition, not an achievement")
        elif non_bracket_runs_italic == 0:
            total_score += 0.2
            print(f"PASS: Component 3 — 0/{non_bracket_runs_total} non-bracketed runs are italic (0.20 pts)")
        else:
            # Partial deduction based on how many were wrongly italicized
            wrong_ratio = non_bracket_runs_italic / non_bracket_runs_total
            comp3_score = 0.2 * max(0, 1.0 - wrong_ratio)
            total_score += comp3_score
            print(f"PARTIAL: Component 3 — {non_bracket_runs_italic}/{non_bracket_runs_total} non-bracketed runs incorrectly italic ({comp3_score:.2f} pts)")

    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
