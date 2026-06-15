"""
Initial Setup: Bilingual document with bracketed translation notes (all regular formatting)
Task ID: writer_frd_031
Domain: libreoffice_writer
"""

import os
import re
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Title --
    title = doc.add_heading("Bilingual Cultural Heritage Guide", level=1)

    # -- Introduction --
    doc.add_heading("Introduction", level=2)
    p1 = doc.add_paragraph()
    p1.add_run(
        "This guide presents an overview of cultural heritage sites across several "
        "European regions, with original terminology preserved alongside English "
        "translations for clarity. The document is intended for international "
        "researchers and heritage conservation professionals."
    )

    # -- Section 1: French Heritage Sites --
    doc.add_heading("Section 1: French Heritage Sites", level=2)

    p2 = doc.add_paragraph()
    p2.add_run(
        "The Chateau de Chambord, located in the Loire Valley, is one of the most "
        "recognizable chateaux in the world. The structure exemplifies French "
        "Renaissance architecture and features a distinctive double-helix staircase "
        "attributed to Leonardo da Vinci. The surrounding estate includes a vast "
        "hunting ground known as the "
    )
    p2.add_run("[domaine de chasse]")  # bracket 1
    p2.add_run(
        " which spans over 5,440 hectares, making it the largest enclosed "
        "park in Europe."
    )

    p3 = doc.add_paragraph()
    p3.add_run(
        "In Provence, visitors can explore the Pont du Gard, an ancient Roman "
        "aqueduct bridge. The locals refer to the surrounding natural area as the "
    )
    p3.add_run("[garrigue mediterraneenne]")  # bracket 2
    p3.add_run(
        ", a type of low scrubland vegetation typical of the Mediterranean basin. "
        "The site has been designated a "
    )
    p3.add_run("[patrimoine mondial de l'UNESCO]")  # bracket 3
    p3.add_run(" since 1985.")

    p4 = doc.add_paragraph()
    p4.add_run(
        "The medieval city of Carcassonne is renowned for its "
    )
    p4.add_run("[cite fortifiee]")  # bracket 4
    p4.add_run(
        ", a fortified city that dates back to the Gallo-Roman period. The "
        "restoration work led by Viollet-le-Duc in the 19th century remains "
        "a subject of scholarly debate regarding "
    )
    p4.add_run("[authenticite architecturale]")  # bracket 5
    p4.add_run(".")

    # -- Section 2: German Heritage Sites --
    doc.add_heading("Section 2: German Heritage Sites", level=2)

    p5 = doc.add_paragraph()
    p5.add_run(
        "The Cologne Cathedral, or Kolner Dom, is a masterpiece of Gothic "
        "architecture. Its construction began in 1248 and was not completed "
        "until 1880. The cathedral houses the Shrine of the Three Kings, "
        "known locally as the "
    )
    p5.add_run("[Dreikonigenschrein]")  # bracket 6
    p5.add_run(
        ", which is considered the largest gilded sarcophagus in the Western world."
    )

    p6 = doc.add_paragraph()
    p6.add_run(
        "In Bavaria, the Neuschwanstein Castle stands as a testament to the "
        "Romantic era. King Ludwig II commissioned the castle as a personal "
    )
    p6.add_run("[Ruckzugsort]")  # bracket 7
    p6.add_run(
        ", a retreat from public life. The interior decorations draw heavily "
        "from the operas of Richard Wagner, featuring scenes from "
    )
    p6.add_run("[Tannhauser und der Sangerkrieg auf Wartburg]")  # bracket 8
    p6.add_run(".")

    # -- Section 3: Italian Heritage Sites --
    doc.add_heading("Section 3: Italian Heritage Sites", level=2)

    p7 = doc.add_paragraph()
    p7.add_run(
        "The Colosseum in Rome, originally known as the Flavian Amphitheatre, "
        "could seat approximately 50,000 spectators. The underground network "
        "of tunnels and chambers beneath the arena floor is called the "
    )
    p7.add_run("[ipogeo]")  # bracket 9
    p7.add_run(
        ", where gladiators and animals were held before events. Recent "
        "archaeological work has revealed a sophisticated "
    )
    p7.add_run("[sistema di elevatori]")  # bracket 10
    p7.add_run(" used to raise combatants to the arena level.")

    p8 = doc.add_paragraph()
    p8.add_run(
        "Venice's Piazza San Marco remains the cultural heart of the city. The "
        "Basilica di San Marco showcases Byzantine architecture and houses "
        "priceless mosaics covering over 8,000 square meters. The periodic "
        "flooding phenomenon known as "
    )
    p8.add_run("[acqua alta]")  # bracket 11
    p8.add_run(
        " continues to threaten the structural integrity of the surrounding "
        "buildings."
    )

    # -- Section 4: Spanish Heritage Sites --
    doc.add_heading("Section 4: Spanish Heritage Sites", level=2)

    p9 = doc.add_paragraph()
    p9.add_run(
        "The Alhambra in Granada represents the pinnacle of Moorish architecture "
        "in Spain. The complex includes the "
    )
    p9.add_run("[Palacios Nazaries]")  # bracket 12
    p9.add_run(
        ", a series of interconnected palaces featuring intricate stucco work "
        "and geometric tile patterns. The Generalife gardens, or "
    )
    p9.add_run("[jardines del Generalife]")  # bracket 13
    p9.add_run(
        ", offer a stunning example of Islamic landscape design with terraced "
        "plantings, fountains, and water channels."
    )

    p10 = doc.add_paragraph()
    p10.add_run(
        "In Barcelona, Antoni Gaudi's Sagrada Familia has been under construction "
        "since 1882. The basilica's design integrates natural forms with religious "
        "symbolism. The main facade, known as the "
    )
    p10.add_run("[fachada de la Natividad]")  # bracket 14
    p10.add_run(
        ", was the first to be completed and depicts scenes from the birth of Christ "
        "with extraordinary sculptural detail."
    )

    # -- Conclusion --
    doc.add_heading("Conclusion", level=2)
    p11 = doc.add_paragraph()
    p11.add_run(
        "The preservation of these cultural heritage sites requires ongoing "
        "international cooperation, adequate funding, and a deep understanding "
        "of local traditions and terminology. This bilingual approach ensures "
        "that the original cultural context is maintained while making the "
        "information accessible to a global audience."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Count brackets for verification
    bracket_count = 0
    for para in doc.paragraphs:
        text = para.text
        bracket_count += len(re.findall(r'\[.+?\]', text))
    print(f'Bracketed segments found: {bracket_count}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
