"""
Initial Setup: 10-page legal document with no headers
Task ID: osworld_writer_header_footer_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_header_footer_004'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup: standard letter size with 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # NO headers or footers — task requires agent to configure them
    # No odd/even page header setting — document starts with standard (no header) state

    # --- Page 1: Title and Introduction ---
    title = doc.add_heading('CONFIDENTIAL SETTLEMENT AGREEMENT', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('Contract Reference: LSA-2025-0847')
    doc.add_paragraph('Effective Date: March 1, 2026')
    doc.add_paragraph('Jurisdiction: State of California, County of Los Angeles')

    doc.add_paragraph(
        'This Settlement Agreement and Release ("Agreement") is entered into as of March 1, 2026, '
        'by and between Meridian Financial Partners LLC, a Delaware limited liability company '
        '("Claimant"), and Harrington & Associates Inc., a California corporation ("Respondent"), '
        'collectively referred to herein as the "Parties."'
    )

    doc.add_paragraph(
        'WHEREAS, a dispute has arisen between the Parties concerning alleged breach of fiduciary '
        'duties, misappropriation of confidential business information, and related claims arising '
        'from the business relationship between the Parties, including but not limited to matters '
        'described in Case No. CV-2025-14823 filed in the Superior Court of the State of California;'
    )

    doc.add_paragraph(
        'WHEREAS, the Parties desire to resolve all such claims, disputes, and controversies between '
        'them without further litigation or arbitration, and to provide for a full and final settlement '
        'of all matters in dispute;'
    )

    doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants and promises contained herein, and '
        'for other good and valuable consideration, the receipt and sufficiency of which are hereby '
        'acknowledged, the Parties agree as follows:'
    )

    doc.add_page_break()

    # --- Page 2: Article I - Definitions ---
    doc.add_heading('ARTICLE I — DEFINITIONS', level=2)

    doc.add_paragraph(
        '1.1 "Confidential Information" means any and all information disclosed by either Party to '
        'the other Party, whether orally, in writing, or by any other means, that is designated as '
        'confidential or that reasonably should be understood to be confidential given the nature of '
        'the information and the circumstances of disclosure, including but not limited to trade '
        'secrets, financial data, customer lists, business plans, proprietary technology, and any '
        'information pertaining to pending litigation or regulatory investigations.'
    )

    doc.add_paragraph(
        '1.2 "Released Claims" means any and all claims, demands, actions, causes of action, suits, '
        'rights, debts, dues, sums of money, accounts, reckonings, covenants, contracts, '
        'controversies, agreements, promises, damages, losses, costs, expenses, and liabilities '
        'of any nature whatsoever, whether known or unknown, suspected or unsuspected, fixed or '
        'contingent, at law or in equity, which either Party has, had, or may have against the other '
        'Party, arising out of or in connection with the matters described in the recitals above.'
    )

    doc.add_paragraph(
        '1.3 "Settlement Amount" means the total sum of Seven Hundred Fifty Thousand Dollars '
        '($750,000.00) USD, to be paid in accordance with the payment schedule set forth in Article III.'
    )

    doc.add_paragraph(
        '1.4 "Effective Date" means March 1, 2026, or such later date as all required signatures '
        'have been obtained from the authorized representatives of both Parties.'
    )

    doc.add_page_break()

    # --- Page 3: Article II - Settlement Terms ---
    doc.add_heading('ARTICLE II — SETTLEMENT TERMS AND CONDITIONS', level=2)

    doc.add_paragraph(
        '2.1 In consideration of the mutual promises and covenants set forth herein, and subject to '
        'the terms and conditions of this Agreement, the Parties agree that the Respondent shall pay '
        'to the Claimant the Settlement Amount as set forth in Article III.'
    )

    doc.add_paragraph(
        '2.2 Upon receipt of the Settlement Amount in full, the Claimant agrees to execute and '
        'deliver to the Respondent a fully executed Stipulation of Dismissal with Prejudice in the '
        'form attached hereto as Exhibit A, dismissing all claims asserted in the above-referenced '
        'litigation with prejudice and without award of attorneys\' fees or costs to either party.'
    )

    doc.add_paragraph(
        '2.3 The Parties acknowledge that this Agreement constitutes a compromise of disputed claims '
        'and shall not be construed as an admission of liability, fault, or wrongdoing by any Party. '
        'The Respondent expressly denies any liability or wrongdoing in connection with the matters '
        'that are the subject of this Agreement.'
    )

    doc.add_paragraph(
        '2.4 Each Party represents and warrants that it has full authority to execute this Agreement '
        'and to perform its obligations hereunder, and that the execution and performance of this '
        'Agreement have been duly authorized by all necessary corporate or organizational action.'
    )

    doc.add_page_break()

    # --- Page 4: Article III - Payment Schedule ---
    doc.add_heading('ARTICLE III — PAYMENT SCHEDULE', level=2)

    doc.add_paragraph(
        '3.1 The Respondent shall pay the Settlement Amount to the Claimant according to the '
        'following installment schedule:'
    )

    payment_items = [
        'First Installment: $250,000.00 due on or before March 15, 2026',
        'Second Installment: $250,000.00 due on or before June 15, 2026',
        'Third Installment: $250,000.00 due on or before September 15, 2026',
    ]
    for item in payment_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '3.2 All payments shall be made by wire transfer to the following account designated by '
        'the Claimant\'s counsel: Meridian Financial Partners LLC Settlement Account, routing '
        'number to be provided by written notice at least five (5) business days prior to each '
        'payment due date.'
    )

    doc.add_paragraph(
        '3.3 In the event any payment is not received by the Claimant within five (5) business '
        'days of the applicable due date, interest shall accrue on the unpaid amount at the rate '
        'of eight percent (8%) per annum from the due date until the date of actual payment. '
        'Additionally, the Claimant shall have the right to reinstate the original litigation '
        'claims upon written notice to the Respondent.'
    )

    doc.add_page_break()

    # --- Page 5: Article IV - Mutual Release ---
    doc.add_heading('ARTICLE IV — MUTUAL RELEASE OF CLAIMS', level=2)

    doc.add_paragraph(
        '4.1 Claimant\'s Release: Effective upon receipt of the Settlement Amount in full, '
        'the Claimant, on behalf of itself and its officers, directors, employees, shareholders, '
        'agents, successors, assigns, heirs, and legal representatives, hereby releases and '
        'forever discharges the Respondent, its officers, directors, employees, shareholders, '
        'agents, successors, and assigns from any and all Released Claims.'
    )

    doc.add_paragraph(
        '4.2 Respondent\'s Release: The Respondent, on behalf of itself and its officers, '
        'directors, employees, shareholders, agents, successors, assigns, heirs, and legal '
        'representatives, hereby releases and forever discharges the Claimant, its officers, '
        'directors, employees, shareholders, agents, successors, and assigns from any and all '
        'Released Claims.'
    )

    doc.add_paragraph(
        '4.3 The Parties expressly waive any and all rights and benefits conferred upon them by '
        'Section 1542 of the California Civil Code, which reads: "A general release does not extend '
        'to claims that the creditor or releasing party does not know or suspect to exist in his or '
        'her favor at the time of executing the release and that, if known by him or her, would have '
        'materially affected his or her settlement with the debtor or released party."'
    )

    doc.add_page_break()

    # --- Page 6: Article V - Confidentiality ---
    doc.add_heading('ARTICLE V — CONFIDENTIALITY OBLIGATIONS', level=2)

    doc.add_paragraph(
        '5.1 Each Party agrees to keep the terms and conditions of this Agreement strictly '
        'confidential and shall not disclose such information to any third party without the prior '
        'written consent of the other Party, except: (i) as required by applicable law, regulation, '
        'or court order; (ii) to legal counsel, accountants, or financial advisors who are bound '
        'by professional obligations of confidentiality; or (iii) in connection with the enforcement '
        'of this Agreement.'
    )

    doc.add_paragraph(
        '5.2 In the event that either Party is required by law, regulation, court order, or '
        'governmental authority to disclose any portion of this Agreement, such Party shall '
        'promptly notify the other Party in writing prior to such disclosure and shall cooperate '
        'in seeking a protective order or other appropriate relief to minimize the scope of such '
        'disclosure.'
    )

    doc.add_paragraph(
        '5.3 The Parties agree that a breach of the confidentiality obligations set forth in this '
        'Article V would cause immediate and irreparable harm for which monetary damages would be '
        'an inadequate remedy, and that the non-breaching Party shall be entitled to seek injunctive '
        'or other equitable relief from any court of competent jurisdiction without the necessity of '
        'proving actual damages or posting bond.'
    )

    doc.add_page_break()

    # --- Page 7: Article VI - Non-Disparagement ---
    doc.add_heading('ARTICLE VI — NON-DISPARAGEMENT AND NON-SOLICITATION', level=2)

    doc.add_paragraph(
        '6.1 Each Party agrees not to make, publish, or communicate, directly or indirectly, any '
        'negative, disparaging, or defamatory statements about the other Party, its officers, '
        'directors, employees, products, services, or business practices to any third party, '
        'including but not limited to media outlets, industry publications, social media platforms, '
        'or professional associations.'
    )

    doc.add_paragraph(
        '6.2 For a period of twenty-four (24) months following the Effective Date, the Respondent '
        'agrees not to directly or indirectly solicit, recruit, or hire any employee, consultant, '
        'or contractor of the Claimant who was employed or engaged by the Claimant at any time '
        'during the twelve (12) months preceding the Effective Date.'
    )

    doc.add_paragraph(
        '6.3 Notwithstanding the foregoing, nothing in this Article shall prohibit either Party '
        'from making truthful statements as required by applicable law or regulation, or from '
        'enforcing its rights under this Agreement.'
    )

    doc.add_page_break()

    # --- Page 8: Article VII - Representations and Warranties ---
    doc.add_heading('ARTICLE VII — REPRESENTATIONS AND WARRANTIES', level=2)

    doc.add_paragraph(
        '7.1 Each Party represents and warrants to the other Party as of the Effective Date that: '
        '(a) it is duly organized, validly existing, and in good standing under the laws of its '
        'jurisdiction of formation; (b) it has full power and authority to execute, deliver, and '
        'perform its obligations under this Agreement; (c) this Agreement constitutes its legal, '
        'valid, and binding obligation, enforceable against it in accordance with its terms; and '
        '(d) the execution, delivery, and performance of this Agreement do not violate any '
        'applicable law, regulation, court order, or contract by which it is bound.'
    )

    doc.add_paragraph(
        '7.2 The Claimant represents and warrants that it is the sole and exclusive owner of the '
        'Released Claims being released herein, has not assigned or transferred any such claims to '
        'any third party, and has full authority to release such claims on behalf of all affiliated '
        'parties as described in Article IV.'
    )

    doc.add_paragraph(
        '7.3 Each Party acknowledges that it has had a full and fair opportunity to consult with '
        'independent legal counsel of its choice prior to executing this Agreement, that it has '
        'read and fully understands the terms and conditions of this Agreement, and that it is '
        'entering into this Agreement freely and voluntarily, without duress or undue influence.'
    )

    doc.add_page_break()

    # --- Page 9: Article VIII - Dispute Resolution ---
    doc.add_heading('ARTICLE VIII — DISPUTE RESOLUTION', level=2)

    doc.add_paragraph(
        '8.1 In the event of any dispute, claim, or controversy arising out of or relating to this '
        'Agreement, or the breach, termination, enforcement, interpretation, or validity thereof, '
        'including the determination of the scope or applicability of this agreement to arbitrate, '
        'the Parties shall first attempt to resolve such dispute through good-faith negotiation '
        'between their respective senior representatives for a period of thirty (30) days following '
        'written notice of the dispute.'
    )

    doc.add_paragraph(
        '8.2 If the dispute cannot be resolved through negotiation within the thirty (30) day period '
        'specified above, the dispute shall be submitted to binding arbitration administered by JAMS '
        'pursuant to its Comprehensive Arbitration Rules and Procedures, before a single arbitrator '
        'who shall be a retired judge or attorney with at least fifteen (15) years of experience in '
        'commercial litigation.'
    )

    doc.add_paragraph(
        '8.3 The arbitration shall be conducted in Los Angeles, California, and the arbitrator\'s '
        'award shall be final and binding upon the Parties. Judgment on the award may be entered in '
        'any court of competent jurisdiction. The Parties agree that the arbitrator shall have '
        'authority to award attorneys\' fees and costs to the prevailing party.'
    )

    doc.add_page_break()

    # --- Page 10: Signature Block ---
    doc.add_heading('ARTICLE IX — GENERAL PROVISIONS', level=2)

    doc.add_paragraph(
        '9.1 Entire Agreement: This Agreement constitutes the entire agreement between the Parties '
        'with respect to the subject matter hereof and supersedes all prior and contemporaneous '
        'negotiations, representations, warranties, agreements, and understandings, whether oral '
        'or written, between the Parties relating to the subject matter of this Agreement.'
    )

    doc.add_paragraph(
        '9.2 Governing Law: This Agreement shall be governed by and construed in accordance with '
        'the laws of the State of California, without regard to any conflict of law principles.'
    )

    doc.add_paragraph(
        '9.3 Counterparts: This Agreement may be executed in any number of counterparts, each of '
        'which shall be deemed an original and all of which together shall constitute one and the '
        'same instrument. Electronic or facsimile signatures shall be deemed original signatures '
        'for all purposes.'
    )

    doc.add_paragraph('\n\nIN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.')

    doc.add_paragraph('MERIDIAN FINANCIAL PARTNERS LLC')
    doc.add_paragraph('By: _______________________________')
    doc.add_paragraph('Name: Katherine R. Vandenberg')
    doc.add_paragraph('Title: Chief Executive Officer')
    doc.add_paragraph('Date: _____________________________')

    doc.add_paragraph('')

    doc.add_paragraph('HARRINGTON & ASSOCIATES INC.')
    doc.add_paragraph('By: _______________________________')
    doc.add_paragraph('Name: Robert T. Harrington III')
    doc.add_paragraph('Title: President and Chief Operating Officer')
    doc.add_paragraph('Date: _____________________________')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
