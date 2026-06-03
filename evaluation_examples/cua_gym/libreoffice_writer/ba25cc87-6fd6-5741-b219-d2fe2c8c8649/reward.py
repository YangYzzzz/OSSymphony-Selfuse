"""
Reward Script: Configure different headers on odd and even pages
Task ID: osworld_writer_header_footer_004
Domain: libreoffice_writer
Scoring:
  - Component 1: evenAndOddHeaders setting enabled (0.2 pts)
  - Component 2: Odd page header has 'Chapter Title' right-aligned (0.3 pts)
  - Component 3: Odd page header contains today's date (0.2 pts)
  - Component 4: Even page header has 'Document Name' left-aligned (0.3 pts)
  Total: 1.0
"""

import os
import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_header_footer_004'


def persist_app_state():
    """Best-effort save via Ctrl+S in case document is open in LibreOffice."""
    try:
        import pyautogui
        import time
        os.environ["DISPLAY"] = ":0"
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.5)
        print("PERSIST: ctrl+s sent")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify that the document has different odd and even page headers configured:
    - Odd pages: 'Chapter Title' right-aligned + today's date
    - Even pages: 'Document Name' left-aligned
    - evenAndOddHeaders setting must be present in document settings

    Returns a float between 0.0 and 1.0.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: evenAndOddHeaders setting enabled (0.2 points)
    # This setting must be present in document settings for odd/even headers to work
    try:
        settings = doc.settings.element
        even_odd_element = settings.find(qn('w:evenAndOddHeaders'))
        if even_odd_element is not None:
            print("PASS: Component 1 — evenAndOddHeaders setting is enabled (0.2 pts)")
            total_score += 0.2
        else:
            print("FAIL: Component 1 — evenAndOddHeaders setting is NOT present in document settings")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Odd page header has 'Chapter Title' and is right-aligned (0.3 points)
    # The default/odd header should contain 'Chapter Title' with RIGHT alignment
    try:
        sec = doc.sections[0]
        odd_hdr = sec.header
        odd_para = odd_hdr.paragraphs[0] if odd_hdr.paragraphs else None

        if odd_para is not None:
            full_text = odd_para.text
            alignment = odd_para.paragraph_format.alignment
            has_chapter_title = "Chapter Title" in full_text
            is_right_aligned = alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT

            if has_chapter_title and is_right_aligned:
                print(f"PASS: Component 2 — Odd header contains 'Chapter Title' right-aligned (0.3 pts)")
                total_score += 0.3
            elif has_chapter_title and not is_right_aligned:
                print(f"FAIL: Component 2 — Odd header has 'Chapter Title' but alignment={alignment} (expected RIGHT)")
            elif is_right_aligned and not has_chapter_title:
                print(f"FAIL: Component 2 — Odd header is right-aligned but missing 'Chapter Title'; text={repr(full_text)}")
            else:
                print(f"FAIL: Component 2 — Odd header missing 'Chapter Title' or not right-aligned; text={repr(full_text)}, alignment={alignment}")
        else:
            print("FAIL: Component 2 — Odd header has no paragraphs")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Odd page header contains today's date (0.2 points)
    # The header should include the current date in 'Month DD, YYYY' format
    try:
        sec = doc.sections[0]
        odd_hdr = sec.header
        odd_para = odd_hdr.paragraphs[0] if odd_hdr.paragraphs else None

        if odd_para is not None:
            full_text = odd_para.text
            today = datetime.date.today()
            # Accept both zero-padded and non-zero-padded day formats
            date_formats = [
                today.strftime("%B %d, %Y"),     # March 06, 2026
                today.strftime("%B %-d, %Y"),    # March 6, 2026
            ]
            date_found = any(fmt in full_text for fmt in date_formats)
            if date_found:
                print(f"PASS: Component 3 — Odd header contains today's date (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 3 — Odd header does not contain today's date; text={repr(full_text)}, expected one of {date_formats}")
        else:
            print("FAIL: Component 3 — Odd header has no paragraphs")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Even page header has 'Document Name' left-aligned (0.3 points)
    # The even header should contain 'Document Name' with LEFT alignment
    try:
        sec = doc.sections[0]
        even_hdr = sec.even_page_header
        even_para = even_hdr.paragraphs[0] if even_hdr.paragraphs else None

        if even_para is not None:
            full_text = even_para.text
            alignment = even_para.paragraph_format.alignment
            has_doc_name = "Document Name" in full_text
            is_left_aligned = alignment in (WD_PARAGRAPH_ALIGNMENT.LEFT, None)
            # None alignment in python-docx means default (left) — check explicit text match
            # For explicit check: alignment == LEFT (0) or alignment is None (inherits left as default)

            if has_doc_name and is_left_aligned:
                print(f"PASS: Component 4 — Even header contains 'Document Name' left-aligned (0.3 pts)")
                total_score += 0.3
            elif has_doc_name and not is_left_aligned:
                print(f"FAIL: Component 4 — Even header has 'Document Name' but alignment={alignment} (expected LEFT or None)")
            elif is_left_aligned and not has_doc_name:
                print(f"FAIL: Component 4 — Even header is left-aligned but missing 'Document Name'; text={repr(full_text)}")
            else:
                print(f"FAIL: Component 4 — Even header missing 'Document Name' or not left-aligned; text={repr(full_text)}, alignment={alignment}")
        else:
            print("FAIL: Component 4 — Even header has no paragraphs")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Attempt to persist any unsaved GUI edits before verification
persist_app_state()

# Run verification against the canonical task file path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
