"""
Initial Setup: Business proposal document with empty title metadata
Task ID: writer_biz_041
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_041'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Ensure document properties title is EMPTY
    doc.core_properties.title = ''
    doc.core_properties.author = 'Meridian Solutions'
    doc.core_properties.subject = 'Partnership Proposal'

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title Page ---
    for _ in range(4):
        doc.add_paragraph('')

    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('MERIDIAN SOLUTIONS')
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run('Partnership Proposal 2025')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x8A)

    doc.add_paragraph('')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run('Prepared for: Apex Industries Ltd.')
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para2 = doc.add_paragraph()
    date_para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run2 = date_para2.add_run('March 2025')
    date_run2.font.size = Pt(12)
    date_run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    conf_run = conf_para.add_run('CONFIDENTIAL')
    conf_run.bold = True
    conf_run.font.size = Pt(10)
    conf_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # --- Table of Contents ---
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Executive Summary', '3'),
        ('2.', 'Company Overview', '4'),
        ('3.', 'Proposed Partnership Structure', '5'),
        ('4.', 'Market Analysis', '6'),
        ('5.', 'Financial Projections', '7'),
        ('6.', 'Implementation Timeline', '8'),
        ('7.', 'Terms and Conditions', '9'),
        ('8.', 'Conclusion', '10'),
    ]
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {title}')
        run.font.size = Pt(12)

    doc.add_page_break()

    # --- 1. Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'Meridian Solutions is pleased to present this partnership proposal to Apex Industries Ltd. '
        'Our organizations share a complementary vision for growth in the enterprise software market, '
        'and we believe a strategic partnership will create significant value for both parties.'
    )
    doc.add_paragraph(
        'Over the past decade, Meridian Solutions has established itself as a leader in cloud-based '
        'business intelligence solutions, serving over 2,400 enterprise clients across North America, '
        'Europe, and Asia-Pacific. Our flagship product, MeridianAI Analytics Suite, processes over '
        '15 billion data points daily for Fortune 500 companies.'
    )
    doc.add_paragraph(
        'This proposal outlines a framework for a mutually beneficial partnership that leverages '
        "Apex Industries' manufacturing expertise and distribution network with Meridian's advanced "
        'analytics capabilities. The projected combined revenue opportunity exceeds $47 million '
        'annually within the first three years of operation.'
    )

    doc.add_page_break()

    # --- 2. Company Overview ---
    doc.add_heading('2. Company Overview', level=1)

    doc.add_heading('2.1 About Meridian Solutions', level=2)
    doc.add_paragraph(
        'Founded in 2012 by Dr. Elena Vasquez and James Chen, Meridian Solutions has grown from a '
        'small startup in Austin, Texas to a global technology firm with 1,850 employees across '
        '12 offices worldwide. Our annual revenue reached $218 million in fiscal year 2024, '
        'representing a 23% year-over-year growth.'
    )

    doc.add_heading('Key Achievements:', level=3)
    achievements = [
        'Named to the Gartner Magic Quadrant for Business Intelligence (2022-2024)',
        'Winner of the TechCrunch Disrupt Enterprise Innovation Award (2023)',
        'ISO 27001 and SOC 2 Type II certified',
        'Net Promoter Score of 72 (industry average: 41)',
        'Strategic partnerships with AWS, Microsoft Azure, and Google Cloud',
    ]
    for item in achievements:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 About Apex Industries', level=2)
    doc.add_paragraph(
        'Apex Industries Ltd., headquartered in Birmingham, United Kingdom, is a leading '
        'manufacturer of industrial automation components with a 45-year track record. '
        'With $1.2 billion in annual revenue and operations spanning 28 countries, Apex '
        'represents an ideal strategic partner for our expansion into the industrial IoT sector.'
    )

    doc.add_page_break()

    # --- 3. Proposed Partnership Structure ---
    doc.add_heading('3. Proposed Partnership Structure', level=1)
    doc.add_paragraph(
        'We propose a three-tiered partnership structure designed to maximize mutual benefit '
        'while maintaining operational independence for both organizations.'
    )

    doc.add_heading('Tier 1: Technology Integration', level=2)
    doc.add_paragraph(
        "Integration of Meridian's analytics platform with Apex's manufacturing execution systems "
        '(MES). This includes real-time data pipelines, predictive maintenance algorithms, and '
        'customized dashboards for plant-level operations managers.'
    )

    doc.add_heading('Tier 2: Joint Go-to-Market', level=2)
    doc.add_paragraph(
        "Combined sales efforts targeting the $12.8 billion industrial analytics market. Apex's "
        "existing relationships with 3,200+ manufacturing clients provides immediate market access "
        "for Meridian's solutions, while our SaaS platform adds recurring revenue streams to "
        "Apex's traditional hardware business."
    )

    doc.add_heading('Tier 3: Co-Innovation Lab', level=2)
    doc.add_paragraph(
        'Establishment of a joint R&D facility in Munich, Germany, with an initial investment '
        'of $8.5 million. The lab will focus on developing next-generation industrial AI solutions '
        'combining edge computing, digital twins, and autonomous decision-making systems.'
    )

    doc.add_page_break()

    # --- 4. Market Analysis ---
    doc.add_heading('4. Market Analysis', level=1)
    doc.add_paragraph(
        'The global industrial analytics market is projected to reach $28.4 billion by 2028, '
        'growing at a CAGR of 15.2% (source: McKinsey Global Industrial Report, 2024). '
        'Key market drivers include:'
    )
    drivers = [
        'Increasing adoption of Industry 4.0 technologies across manufacturing sectors',
        'Growing demand for predictive maintenance solutions (projected $18.6B by 2027)',
        'Regulatory requirements for real-time emissions and quality monitoring',
        'Labor shortages driving automation investment in developed economies',
        'Edge computing advancements enabling real-time analytics at the factory floor',
    ]
    for d in drivers:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph(
        'Our competitive analysis identifies a significant gap in the market for integrated '
        'analytics-manufacturing solutions. Current vendors offer either analytics platforms '
        '(Splunk, Palantir, Databricks) or manufacturing automation (Siemens, Rockwell, ABB) '
        'but not an end-to-end solution that bridges both domains.'
    )

    doc.add_page_break()

    # --- 5. Financial Projections ---
    doc.add_heading('5. Financial Projections', level=1)
    doc.add_paragraph(
        'Below is a summary of the projected financial performance for the partnership '
        'over the first five years:'
    )

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['Metric', 'Year 1', 'Year 3', 'Year 5']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Combined Revenue', '$12.4M', '$47.2M', '$89.6M'],
        ['Cost Savings (Operational)', '$2.1M', '$8.7M', '$14.3M'],
        ['Joint Client Acquisitions', '45', '280', '620'],
        ['R&D Investment', '$8.5M', '$5.2M', '$4.8M'],
        ['Net Profit Margin', '8.2%', '18.5%', '24.1%'],
        ['ROI', '14%', '62%', '118%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'These projections are based on conservative estimates derived from comparable '
        'technology-manufacturing partnerships in the market. The break-even point for '
        'initial investment is projected at 26 months from partnership commencement.'
    )

    doc.add_page_break()

    # --- 6. Implementation Timeline ---
    doc.add_heading('6. Implementation Timeline', level=1)

    timeline_table = doc.add_table(rows=6, cols=3)
    timeline_table.style = 'Table Grid'
    t_headers = ['Phase', 'Timeline', 'Key Milestones']
    for i, h in enumerate(t_headers):
        cell = timeline_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    t_data = [
        ['Phase 1: Foundation', 'Q2-Q3 2025', 'Legal agreements, team formation, API integration planning'],
        ['Phase 2: Integration', 'Q4 2025 - Q1 2026', 'Platform integration, pilot program with 10 clients'],
        ['Phase 3: Launch', 'Q2 2026', 'Joint product launch, marketing campaign, sales enablement'],
        ['Phase 4: Scale', 'Q3-Q4 2026', 'Geographic expansion, 100+ client onboarding'],
        ['Phase 5: Optimize', '2027+', 'Performance optimization, co-innovation lab launch'],
    ]
    for r, row_data in enumerate(t_data, 1):
        for c, val in enumerate(row_data):
            timeline_table.cell(r, c).text = val

    doc.add_page_break()

    # --- 7. Terms and Conditions ---
    doc.add_heading('7. Terms and Conditions', level=1)
    doc.add_paragraph(
        'The following terms outline the proposed framework for the partnership. '
        'Final terms will be subject to mutual agreement and legal review.'
    )

    terms = [
        'Revenue Sharing: 60/40 split (Meridian/Apex) for analytics-driven revenue; '
        '40/60 split for hardware-driven revenue.',
        'Intellectual Property: Joint IP for co-developed solutions; pre-existing IP '
        'remains with the originating party.',
        'Exclusivity: Non-exclusive partnership with a 12-month right of first refusal '
        'for competing partnerships in the industrial analytics sector.',
        'Term: Initial 5-year agreement with automatic 2-year renewal periods.',
        'Governance: Joint steering committee with equal representation, meeting quarterly.',
        'Termination: Either party may terminate with 180 days written notice after the '
        'initial 24-month commitment period.',
        'Data Protection: Full compliance with GDPR, CCPA, and applicable local data '
        'protection regulations.',
        'Liability: Capped at 12 months of fees paid under the partnership agreement.',
    ]
    for i, term in enumerate(terms, 1):
        doc.add_paragraph(f'{i}. {term}')

    doc.add_page_break()

    # --- 8. Conclusion ---
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph(
        'Meridian Solutions is confident that this partnership with Apex Industries will '
        'create transformative value for both organizations. By combining our respective '
        'strengths in analytics and manufacturing, we can establish a market-leading position '
        'in the rapidly growing industrial intelligence sector.'
    )
    doc.add_paragraph(
        'We look forward to discussing this proposal in detail and exploring the significant '
        'opportunities that lie ahead. Our team is available for a comprehensive presentation '
        'and Q&A session at your earliest convenience.'
    )

    doc.add_paragraph('')
    contact_para = doc.add_paragraph()
    contact_run = contact_para.add_run('Contact Information:')
    contact_run.bold = True
    doc.add_paragraph('Dr. Elena Vasquez, CEO')
    doc.add_paragraph('elena.vasquez@meridiansolutions.com')
    doc.add_paragraph('+1 (512) 555-0142')
    doc.add_paragraph('')
    doc.add_paragraph('James Chen, VP of Strategic Partnerships')
    doc.add_paragraph('james.chen@meridiansolutions.com')
    doc.add_paragraph('+1 (512) 555-0198')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
