"""
Initial Setup: Figure numbering in master document with chapter-based numbering
Task ID: writer_rm_073
Domain: libreoffice_writer

Creates a document with 5 chapters, each containing figures with inconsistent
caption numbering. The agent must fix them to use 'Figure X.Y' format.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_073'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_figure_placeholder(doc, width_inches=4.0, height_inches=2.5, label="Figure"):
    """Add a gray rectangle placeholder representing a figure."""
    # We'll add a simple text-based placeholder since we can't easily generate images
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(f"[{label} - Image Placeholder]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True
    return para


def add_caption(doc, caption_text, style="Normal"):
    """Add a figure caption paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(caption_text)
    run.font.size = Pt(10)
    run.font.italic = True
    return para


def create_initial():
    doc = Document()

    # ---- Title Page ----
    title = doc.add_heading("Advanced Data Analytics: A Comprehensive Textbook", level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    run = subtitle.add_run("Second Edition")
    run.font.size = Pt(16)
    run.font.italic = True

    authors = doc.add_paragraph()
    authors.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors.paragraph_format.space_before = Pt(36)
    run = authors.add_run("Dr. Elena Rodriguez & Prof. James Mitchell")
    run.font.size = Pt(14)

    publisher = doc.add_paragraph()
    publisher.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    publisher.paragraph_format.space_before = Pt(48)
    run = publisher.add_run("Academic Press International, 2025")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()

    # ===========================================================
    # CHAPTER 1: Introduction to Data Analytics
    # ===========================================================
    doc.add_heading("Chapter 1: Introduction to Data Analytics", level=1)

    doc.add_paragraph(
        "Data analytics has transformed the way organizations make decisions. "
        "In the modern business landscape, the ability to collect, process, and "
        "interpret large volumes of data has become a critical competitive advantage. "
        "This chapter provides an overview of the fundamental concepts and methodologies "
        "that form the backbone of contemporary data analytics practice."
    )

    doc.add_heading("1.1 The Evolution of Data-Driven Decision Making", level=2)
    doc.add_paragraph(
        "The journey from intuition-based to data-driven decision making spans several "
        "decades. Early statistical methods developed in the 1950s laid the groundwork "
        "for what we now recognize as modern analytics. The advent of computing power "
        "in the 1980s and 1990s accelerated the adoption of quantitative approaches "
        "across industries."
    )

    # Figure 1 - INCONSISTENT: uses "Fig. 1" instead of "Figure 1.1"
    add_figure_placeholder(doc, label="Timeline of analytics evolution")
    add_caption(doc, "Fig. 1: Timeline of Analytics Evolution from 1950s to Present")

    doc.add_paragraph(
        "As shown in the timeline above, each decade brought significant advancements. "
        "The 2000s saw the emergence of big data platforms, while the 2010s introduced "
        "machine learning at scale. Today, AI-powered analytics represents the cutting edge."
    )

    doc.add_heading("1.2 Key Concepts and Terminology", level=2)
    doc.add_paragraph(
        "Understanding the core terminology is essential for any practitioner. "
        "Descriptive analytics tells us what happened, diagnostic analytics explains "
        "why it happened, predictive analytics forecasts what might happen, and "
        "prescriptive analytics recommends actions to take."
    )

    # Figure 2 - INCONSISTENT: uses "Figure 2" (no chapter prefix)
    add_figure_placeholder(doc, label="Analytics maturity model diagram")
    add_caption(doc, "Figure 2: The Four Types of Analytics Maturity Model")

    doc.add_heading("1.3 The Data Analytics Lifecycle", level=2)
    doc.add_paragraph(
        "Every analytics project follows a structured lifecycle. From problem definition "
        "through data collection, cleaning, analysis, and visualization, each phase "
        "requires specific skills and tools. The iterative nature of this process means "
        "that analysts often revisit earlier stages as new insights emerge."
    )

    # Figure 3 - INCONSISTENT: uses "Figure 1-3" (hyphen instead of dot)
    add_figure_placeholder(doc, label="Data analytics lifecycle flowchart")
    add_caption(doc, "Figure 1-3: The Complete Data Analytics Lifecycle")

    doc.add_page_break()

    # ===========================================================
    # CHAPTER 2: Statistical Foundations
    # ===========================================================
    doc.add_heading("Chapter 2: Statistical Foundations", level=1)

    doc.add_paragraph(
        "A solid grounding in statistics is indispensable for meaningful data analysis. "
        "This chapter reviews the essential statistical concepts that every data analyst "
        "must master, from probability distributions to hypothesis testing frameworks."
    )

    doc.add_heading("2.1 Probability Distributions", level=2)
    doc.add_paragraph(
        "Probability distributions describe how values of a random variable are spread. "
        "The normal distribution, also known as the Gaussian distribution, is perhaps "
        "the most important in statistics due to the Central Limit Theorem. Other key "
        "distributions include the binomial, Poisson, and exponential distributions."
    )

    # Figure 1 of Ch2 - INCONSISTENT: uses "Figure 4" (sequential, ignoring chapters)
    add_figure_placeholder(doc, label="Normal distribution bell curve")
    add_caption(doc, "Figure 4: Standard Normal Distribution with Mean and Standard Deviation")

    doc.add_heading("2.2 Hypothesis Testing", level=2)
    doc.add_paragraph(
        "Hypothesis testing provides a formal framework for making statistical decisions. "
        "The null hypothesis represents the status quo, while the alternative hypothesis "
        "represents the claim we want to test. P-values, significance levels, and "
        "confidence intervals are the tools of this framework."
    )

    # Figure 2 of Ch2 - INCONSISTENT: uses "Fig 5" (abbreviated, no period)
    add_figure_placeholder(doc, label="Hypothesis testing decision tree")
    add_caption(doc, "Fig 5: Decision Framework for Hypothesis Testing")

    doc.add_heading("2.3 Regression Analysis", level=2)
    doc.add_paragraph(
        "Regression analysis examines the relationship between dependent and independent "
        "variables. Simple linear regression models the relationship between two variables, "
        "while multiple regression extends this to several predictors. Understanding "
        "assumptions such as linearity, normality of residuals, and homoscedasticity "
        "is critical for valid inference."
    )

    # Figure 3 of Ch2 - INCONSISTENT: uses "Figure II-3" (Roman numeral chapter)
    add_figure_placeholder(doc, label="Scatter plot with regression line")
    add_caption(doc, "Figure II-3: Linear Regression Fit with Residual Plot")

    doc.add_heading("2.4 Bayesian Statistics", level=2)
    doc.add_paragraph(
        "Bayesian methods offer an alternative to frequentist statistics by incorporating "
        "prior knowledge into the analysis. Bayes' theorem provides the mathematical "
        "foundation for updating beliefs in light of new evidence. This approach is "
        "particularly valuable when dealing with small sample sizes or incorporating "
        "expert knowledge."
    )

    # Figure 4 of Ch2 - INCONSISTENT: uses "Diagram 7"
    add_figure_placeholder(doc, label="Bayesian updating process")
    add_caption(doc, "Diagram 7: Bayesian Updating Process with Prior and Posterior Distributions")

    doc.add_page_break()

    # ===========================================================
    # CHAPTER 3: Data Visualization Techniques
    # ===========================================================
    doc.add_heading("Chapter 3: Data Visualization Techniques", level=1)

    doc.add_paragraph(
        "Effective data visualization is both an art and a science. The ability to "
        "present complex information in an intuitive visual format is crucial for "
        "communicating insights to stakeholders. This chapter covers fundamental "
        "visualization principles and advanced charting techniques."
    )

    doc.add_heading("3.1 Principles of Effective Visualization", level=2)
    doc.add_paragraph(
        "Edward Tufte's principles of data-ink ratio and chart junk avoidance remain "
        "foundational. A good visualization should maximize the data-ink ratio while "
        "minimizing non-essential visual elements. Color choice, typography, and "
        "layout all play important roles in conveying information clearly."
    )

    # Figure 1 of Ch3 - INCONSISTENT: uses "Figure 8" (continuing sequential)
    add_figure_placeholder(doc, label="Good vs bad visualization comparison")
    add_caption(doc, "Figure 8: Comparison of Effective vs Ineffective Chart Design")

    doc.add_heading("3.2 Chart Selection Guide", level=2)
    doc.add_paragraph(
        "Choosing the right chart type depends on the nature of the data and the "
        "message you want to convey. Bar charts work well for categorical comparisons, "
        "line charts for trends over time, scatter plots for relationships between "
        "variables, and pie charts (used sparingly) for part-to-whole relationships."
    )

    # Figure 2 of Ch3 - INCONSISTENT: uses "Fig. 3.2" (period after Fig)
    add_figure_placeholder(doc, label="Chart type decision matrix")
    add_caption(doc, "Fig. 3.2: Chart Type Selection Decision Matrix")

    doc.add_heading("3.3 Interactive Dashboards", level=2)
    doc.add_paragraph(
        "Modern analytics platforms enable the creation of interactive dashboards "
        "that allow users to explore data dynamically. Tools like Tableau, Power BI, "
        "and custom web-based solutions using D3.js provide flexible visualization "
        "capabilities. Dashboard design should follow the principle of progressive "
        "disclosure, showing high-level summaries first with drill-down options."
    )

    # Figure 3 of Ch3 - INCONSISTENT: uses "FIGURE 10" (all caps)
    add_figure_placeholder(doc, label="Interactive dashboard mockup")
    add_caption(doc, "FIGURE 10: Sample Interactive Analytics Dashboard Layout")

    doc.add_heading("3.4 Geospatial Visualization", level=2)
    doc.add_paragraph(
        "Geospatial data requires specialized visualization approaches. Choropleth "
        "maps, heat maps, and point maps each serve different purposes. Geographic "
        "Information Systems (GIS) tools like QGIS and ArcGIS provide sophisticated "
        "mapping capabilities, while web libraries such as Leaflet and Mapbox enable "
        "interactive map creation."
    )

    # Figure 4 of Ch3 - INCONSISTENT: uses "figure 11" (lowercase)
    add_figure_placeholder(doc, label="Choropleth map example")
    add_caption(doc, "figure 11: Choropleth Map Showing Regional Sales Distribution")

    doc.add_page_break()

    # ===========================================================
    # CHAPTER 4: Machine Learning Fundamentals
    # ===========================================================
    doc.add_heading("Chapter 4: Machine Learning Fundamentals", level=1)

    doc.add_paragraph(
        "Machine learning represents a paradigm shift in how we approach data analysis. "
        "Rather than explicitly programming rules, machine learning algorithms learn "
        "patterns from data. This chapter introduces the core concepts, algorithms, "
        "and best practices for applying machine learning to real-world problems."
    )

    doc.add_heading("4.1 Supervised Learning", level=2)
    doc.add_paragraph(
        "Supervised learning algorithms learn from labeled training data. Classification "
        "tasks assign categories (spam vs. not spam), while regression tasks predict "
        "continuous values (house prices). Common algorithms include decision trees, "
        "random forests, support vector machines, and neural networks."
    )

    # Figure 1 of Ch4 - INCONSISTENT: uses "Figure A" (letter instead of number)
    add_figure_placeholder(doc, label="Supervised learning workflow")
    add_caption(doc, "Figure A: Supervised Learning Training and Prediction Pipeline")

    doc.add_heading("4.2 Unsupervised Learning", level=2)
    doc.add_paragraph(
        "Unsupervised learning discovers hidden patterns in unlabeled data. Clustering "
        "algorithms like K-means and DBSCAN group similar data points together. "
        "Dimensionality reduction techniques such as PCA and t-SNE help visualize "
        "high-dimensional data in two or three dimensions."
    )

    # Figure 2 of Ch4 - INCONSISTENT: uses "Figure #13" (hash symbol)
    add_figure_placeholder(doc, label="K-means clustering visualization")
    add_caption(doc, "Figure #13: K-Means Clustering with Three Distinct Groups")

    doc.add_heading("4.3 Model Evaluation", level=2)
    doc.add_paragraph(
        "Evaluating model performance requires appropriate metrics and validation "
        "strategies. For classification, metrics include accuracy, precision, recall, "
        "F1-score, and AUC-ROC. Cross-validation techniques such as k-fold and "
        "stratified sampling help assess how well a model generalizes to unseen data."
    )

    # Figure 3 of Ch4 - INCONSISTENT: uses "Fig (14)" (parentheses)
    add_figure_placeholder(doc, label="ROC curve comparison")
    add_caption(doc, "Fig (14): ROC Curves Comparing Three Classification Models")

    doc.add_heading("4.4 Feature Engineering", level=2)
    doc.add_paragraph(
        "Feature engineering is often the most impactful step in machine learning. "
        "Creating informative features from raw data can dramatically improve model "
        "performance. Techniques include one-hot encoding, binning, polynomial features, "
        "and domain-specific transformations. Feature selection methods help identify "
        "the most relevant predictors."
    )

    # Figure 4 of Ch4 - INCONSISTENT: uses "Exhibit 15"
    add_figure_placeholder(doc, label="Feature importance ranking chart")
    add_caption(doc, "Exhibit 15: Feature Importance Rankings from Random Forest Model")

    # Figure 5 of Ch4 - INCONSISTENT: uses "Figure 4/5"
    add_figure_placeholder(doc, label="Feature engineering pipeline")
    add_caption(doc, "Figure 4/5: End-to-End Feature Engineering Pipeline")

    doc.add_page_break()

    # ===========================================================
    # CHAPTER 5: Big Data Technologies
    # ===========================================================
    doc.add_heading("Chapter 5: Big Data Technologies", level=1)

    doc.add_paragraph(
        "The explosive growth of data has necessitated new technologies for storage, "
        "processing, and analysis. This chapter explores the ecosystem of big data "
        "tools and frameworks that enable organizations to work with datasets too "
        "large for traditional database systems."
    )

    doc.add_heading("5.1 Distributed Computing Frameworks", level=2)
    doc.add_paragraph(
        "Apache Hadoop pioneered distributed data processing with its MapReduce "
        "paradigm. Apache Spark subsequently revolutionized the field by offering "
        "in-memory processing that can be orders of magnitude faster. Understanding "
        "the architecture of these systems is essential for building scalable "
        "analytics pipelines."
    )

    # Figure 1 of Ch5 - INCONSISTENT: uses "Figure 16" (sequential)
    add_figure_placeholder(doc, label="Hadoop vs Spark architecture comparison")
    add_caption(doc, "Figure 16: Architecture Comparison of Hadoop MapReduce and Apache Spark")

    doc.add_heading("5.2 Cloud Data Platforms", level=2)
    doc.add_paragraph(
        "Cloud platforms have democratized access to big data infrastructure. "
        "Amazon Web Services (AWS), Google Cloud Platform (GCP), and Microsoft Azure "
        "each offer comprehensive suites of data services. Key services include "
        "data warehouses (BigQuery, Redshift), streaming platforms (Kinesis, Pub/Sub), "
        "and managed machine learning environments."
    )

    # Figure 2 of Ch5 - INCONSISTENT: uses "Fig V-2" (Roman numeral, no period)
    add_figure_placeholder(doc, label="Cloud platform service comparison")
    add_caption(doc, "Fig V-2: Comparison of Major Cloud Data Platform Services")

    doc.add_heading("5.3 Real-Time Data Processing", level=2)
    doc.add_paragraph(
        "Real-time analytics requires specialized architectures. Apache Kafka provides "
        "a distributed messaging system for streaming data. Apache Flink and Spark "
        "Structured Streaming enable real-time computation on data streams. The Lambda "
        "and Kappa architectures provide design patterns for systems that need both "
        "batch and real-time processing capabilities."
    )

    # Figure 3 of Ch5 - INCONSISTENT: uses "Figure [18]" (brackets)
    add_figure_placeholder(doc, label="Lambda architecture diagram")
    add_caption(doc, "Figure [18]: Lambda Architecture for Combined Batch and Real-Time Processing")

    doc.add_heading("5.4 Data Governance and Ethics", level=2)
    doc.add_paragraph(
        "As organizations collect and process ever-larger volumes of data, governance "
        "and ethical considerations become paramount. Data lineage tracking, access "
        "controls, privacy regulations (GDPR, CCPA), and algorithmic fairness are "
        "all critical aspects of responsible data management. Organizations must "
        "establish clear data governance frameworks that balance innovation with "
        "compliance and ethical responsibility."
    )

    # Figure 4 of Ch5 - INCONSISTENT: uses "figure nineteen"
    add_figure_placeholder(doc, label="Data governance framework")
    add_caption(doc, "figure nineteen: Comprehensive Data Governance Framework Overview")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
