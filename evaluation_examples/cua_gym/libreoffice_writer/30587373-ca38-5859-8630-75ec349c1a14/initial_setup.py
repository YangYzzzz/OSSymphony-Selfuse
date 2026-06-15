"""
Initial Setup: Legal brief document with headings but no automatic numbering
Task ID: writer_fp_048
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('MEMORANDUM OF LAW IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Case No. 2025-CV-04817 | Meridian Corp. v. Ashford Holdings LLC')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    # ---- Heading 1: #1 ----
    doc.add_heading('Jurisdiction', level=1)
    p = doc.add_paragraph(
        'This Court has subject matter jurisdiction over this action pursuant to 28 U.S.C. '
        '\u00a7 1332, as the matter involves citizens of different states and the amount in '
        'controversy exceeds $75,000, exclusive of interest and costs.'
    )
    p.paragraph_format.space_after = Pt(6)

    # H2 under Jurisdiction
    doc.add_heading('Personal Jurisdiction', level=2)
    doc.add_paragraph(
        'Defendant Ashford Holdings LLC is a Delaware limited liability company with its '
        'principal place of business in Harris County, Texas. Defendant has engaged in '
        'continuous and systematic business activities within this judicial district, '
        'establishing general jurisdiction under the International Shoe framework.'
    )

    doc.add_heading('Subject Matter Jurisdiction', level=2)
    doc.add_paragraph(
        'Complete diversity exists between the parties. Plaintiff Meridian Corp. is '
        'incorporated in California with its principal office in San Francisco. The amount '
        'in controversy, as established by the complaint and supporting declarations, '
        'exceeds the statutory threshold of $75,000.'
    )

    doc.add_heading('Venue', level=2)
    doc.add_paragraph(
        'Venue is proper in this district pursuant to 28 U.S.C. \u00a7 1391(b)(2), as a '
        'substantial part of the events giving rise to the claims occurred within this district.'
    )

    # ---- Heading 1: #2 ----
    doc.add_heading('Statement of Facts', level=1)
    doc.add_paragraph(
        'On March 15, 2024, Plaintiff Meridian Corp. and Defendant Ashford Holdings LLC '
        'entered into a Master Services Agreement ("MSA") governing the development and '
        'deployment of an enterprise resource planning system for Defendant\'s North American '
        'operations. The total contract value was $4.2 million, payable in four quarterly '
        'installments upon achievement of specified milestones.'
    )

    doc.add_heading('Contract Formation', level=2)
    doc.add_paragraph(
        'The MSA was executed following six months of negotiation between the parties. '
        'Sarah Chen, Vice President of Business Development at Meridian, and Marcus Johnson, '
        'Chief Technology Officer at Ashford Holdings, served as the primary negotiators. '
        'The agreement incorporated detailed specifications set forth in Exhibit A.'
    )

    # H3 under Contract Formation
    doc.add_heading('Negotiation History', level=3)
    doc.add_paragraph(
        'Initial discussions commenced in September 2023, when Ashford issued a Request for '
        'Proposal to twelve qualified vendors. Meridian submitted its response on October 12, '
        '2023, and was selected as the preferred vendor on November 28, 2023.'
    )

    doc.add_heading('Material Terms', level=3)
    doc.add_paragraph(
        'The MSA contained the following material provisions: (a) a fixed-price structure of '
        '$4.2 million; (b) quarterly milestone payments of $1.05 million each; (c) a '
        'completion deadline of March 15, 2025; and (d) a limitation of liability capped '
        'at the total contract price.'
    )

    doc.add_heading('Performance Period', level=2)
    doc.add_paragraph(
        'Meridian commenced performance on April 1, 2024, deploying a team of fourteen '
        'engineers and three project managers to Ashford\'s Houston facility. The first '
        'milestone \u2014 completion of system architecture and database design \u2014 was '
        'achieved on schedule on June 30, 2024.'
    )

    doc.add_heading('Breach Events', level=2)
    doc.add_paragraph(
        'Despite timely completion of the first two milestones, Defendant failed to remit '
        'the second quarterly payment of $1.05 million due on September 30, 2024. Defendant '
        'further repudiated the agreement on November 15, 2024, by terminating Meridian\'s '
        'access to its facilities and systems without cause or notice.'
    )

    # H3 under Breach Events
    doc.add_heading('Notice of Default', level=3)
    doc.add_paragraph(
        'On October 15, 2024, Meridian delivered written notice of default via certified '
        'mail to Ashford\'s registered agent, as required by Section 12.3 of the MSA. '
        'The notice specified the outstanding payment and provided thirty days to cure.'
    )

    # ---- Heading 1: #3 ----
    doc.add_heading('Applicable Law', level=1)
    doc.add_paragraph(
        'The substantive claims in this action are governed by Texas law, as expressly '
        'provided in Section 15.1 of the MSA. Texas follows the Restatement (Second) of '
        'Contracts for interpretation of contractual obligations and remedies.'
    )

    doc.add_heading('Contract Interpretation Standards', level=2)
    doc.add_paragraph(
        'Under Texas law, the primary goal of contract interpretation is to ascertain and '
        'give effect to the parties\' intent as expressed in the written instrument. '
        'Coker v. Coker, 650 S.W.2d 391, 393 (Tex. 1983). When the language of a '
        'contract is unambiguous, the court must enforce it as written.'
    )

    doc.add_heading('Parol Evidence Rule', level=2)
    doc.add_paragraph(
        'The MSA contains a comprehensive integration clause at Section 16.1, establishing '
        'that the written agreement constitutes the entire understanding between the parties. '
        'Under Texas\'s parol evidence rule, extrinsic evidence of prior or contemporaneous '
        'agreements may not be used to contradict the terms of the integrated agreement.'
    )

    # H3 under Parol Evidence
    doc.add_heading('Exceptions to Parol Evidence', level=3)
    doc.add_paragraph(
        'The exceptions recognized under Texas law \u2014 fraud, duress, and mutual mistake \u2014 '
        'are not applicable here. Defendant has not alleged any of these exceptions in its '
        'answer or affirmative defenses.'
    )

    doc.add_heading('Damages Framework', level=2)
    doc.add_paragraph(
        'Texas applies the expectation interest measure of damages for breach of contract, '
        'seeking to place the non-breaching party in the position it would have occupied '
        'had the contract been fully performed. Stewart v. Basey, 245 S.W.2d 484, 486 '
        '(Tex. 1952).'
    )

    # ---- Heading 1: #4 ----
    doc.add_heading('Legal Argument', level=1)
    doc.add_paragraph(
        'Plaintiff is entitled to summary judgment on its breach of contract claim because '
        'there are no genuine disputes of material fact, and the undisputed record '
        'establishes every element of the claim as a matter of law.'
    )

    doc.add_heading('Elements of Breach of Contract', level=2)
    doc.add_paragraph(
        'Under Texas law, a plaintiff asserting breach of contract must establish: '
        '(1) the existence of a valid contract; (2) performance or tendered performance '
        'by the plaintiff; (3) breach by the defendant; and (4) damages resulting from '
        'the breach. Valero Mktg. & Supply Co. v. Kalama Int\'l, 51 S.W.3d 345, 351 '
        '(Tex. App.\u2014Houston [1st Dist.] 2001).'
    )

    # H3 under Elements
    doc.add_heading('Valid Contract', level=3)
    doc.add_paragraph(
        'The existence of a valid, enforceable contract is not in dispute. Both parties '
        'signed the MSA on March 15, 2024, and have acknowledged its validity in their '
        'respective pleadings. The contract satisfies all requisite elements: offer, '
        'acceptance, consideration, and mutual assent.'
    )

    doc.add_heading('Plaintiff Performance', level=3)
    doc.add_paragraph(
        'Meridian performed its obligations under the MSA through the completion of '
        'Milestones 1 and 2, as certified by Defendant\'s own project management office '
        'in acceptance letters dated June 30 and September 15, 2024, respectively.'
    )

    doc.add_heading('Defendant Breach', level=3)
    doc.add_paragraph(
        'Ashford\'s failure to remit the second quarterly payment and its subsequent '
        'wrongful termination of the agreement constitute material breaches of the MSA. '
        'The payment obligation was unconditional upon milestone acceptance, and the '
        'termination was effected without the required thirty-day notice and cure period.'
    )

    doc.add_heading('Summary Judgment Standard', level=2)
    doc.add_paragraph(
        'Summary judgment is proper when the movant demonstrates that there is no genuine '
        'dispute as to any material fact and the movant is entitled to judgment as a matter '
        'of law. Fed. R. Civ. P. 56(a). The court must view the evidence in the light '
        'most favorable to the non-moving party.'
    )

    doc.add_heading('Undisputed Material Facts', level=2)
    doc.add_paragraph(
        'The following facts are established by sworn declarations, authenticated documents, '
        'and Defendant\'s own admissions in discovery responses, and are not subject to '
        'genuine dispute: the MSA was validly executed; Meridian completed Milestones 1 '
        'and 2; Ashford accepted both milestones; and Ashford failed to pay $1.05 million '
        'due on September 30, 2024.'
    )

    # ---- Heading 1: #5 ----
    doc.add_heading('Damages and Relief', level=1)
    doc.add_paragraph(
        'Plaintiff seeks compensatory damages in the amount of $2,310,000, comprising '
        'the unpaid second installment ($1,050,000), lost profit on the remaining two '
        'installments ($840,000), and consequential damages for costs incurred in '
        'reassigning personnel ($420,000).'
    )

    doc.add_heading('Direct Damages', level=2)
    doc.add_paragraph(
        'The unpaid second milestone payment of $1,050,000 constitutes direct damages '
        'owed under the express terms of the MSA. This amount was due and payable on '
        'September 30, 2024, upon Ashford\'s written acceptance of Milestone 2.'
    )

    doc.add_heading('Consequential Damages', level=2)
    doc.add_paragraph(
        'As a direct and foreseeable consequence of Defendant\'s breach, Meridian incurred '
        '$420,000 in costs associated with reassigning its seventeen-person project team '
        'to other engagements. These costs were within the contemplation of the parties '
        'at the time of contracting, as evidenced by the MSA\'s staffing provisions.'
    )

    doc.add_heading('Prejudgment Interest', level=2)
    doc.add_paragraph(
        'Plaintiff is entitled to prejudgment interest at the statutory rate of 5% per '
        'annum from September 30, 2024, the date the first unpaid installment became due, '
        'through the date of judgment. Tex. Fin. Code \u00a7 304.003.'
    )

    # H3 under Prejudgment Interest
    doc.add_heading('Calculation Methodology', level=3)
    doc.add_paragraph(
        'The statutory prejudgment interest rate is applied on a simple interest basis from '
        'the date each payment became due. For the second installment of $1,050,000, interest '
        'accrues from September 30, 2024. The daily rate of $143.84 yields total prejudgment '
        'interest of approximately $52,500 through the anticipated trial date.'
    )

    # Set default font for body paragraphs
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
