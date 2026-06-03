"""
Initial Setup: Legal brief with mixed tab indentation
Task ID: writer_frd_011
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_011'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title Block ---
    # Tab count: 3 tabs in the title block
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run("IN THE UNITED STATES DISTRICT COURT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("FOR THE NORTHERN DISTRICT OF CALIFORNIA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()  # blank line

    # Case caption with tabs (3 tabs)
    caption = doc.add_paragraph()
    run = caption.add_run("MERIDIAN TECHNOLOGIES, INC.,\tPlaintiff,")  # tab 1
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    caption2 = doc.add_paragraph()
    run = caption2.add_run("\tv.\tCase No. 24-CV-03847-RMW")  # tabs 2,3
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    caption3 = doc.add_paragraph()
    run = caption3.add_run("APEX SOLUTIONS GROUP, LLC,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    caption4 = doc.add_paragraph()
    run = caption4.add_run("Defendant.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Separator
    sep = doc.add_paragraph()
    sep.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sep.add_run("_" * 60)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Document title
    doc_title = doc.add_paragraph()
    doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = doc_title.add_run("PLAINTIFF'S MEMORANDUM OF LAW IN SUPPORT OF")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    doc_title2 = doc.add_paragraph()
    doc_title2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = doc_title2.add_run("MOTION FOR SUMMARY JUDGMENT")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    # Separator
    sep2 = doc.add_paragraph()
    sep2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = sep2.add_run("_" * 60)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # --- I. INTRODUCTION --- (4 tabs)
    heading1 = doc.add_paragraph()
    run = heading1.add_run("I.\tINTRODUCTION")  # tab 4
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    intro1 = doc.add_paragraph()
    run = intro1.add_run(
        "\tPlaintiff Meridian Technologies, Inc. (\"Meridian\") respectfully submits this "  # tab 5
        "memorandum of law in support of its motion for summary judgment against Defendant "
        "Apex Solutions Group, LLC (\"Apex\"). The undisputed facts demonstrate that Apex "
        "willfully misappropriated Meridian's proprietary source code and trade secrets in "
        "violation of the Defend Trade Secrets Act, 18 U.S.C. § 1836 et seq."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    intro2 = doc.add_paragraph()
    run = intro2.add_run(
        "\tAs set forth below, the evidence conclusively establishes that: (a) Meridian's "  # tab 6
        "proprietary algorithms constitute protectable trade secrets; (b) Apex acquired these "
        "trade secrets through improper means; and (c) Meridian has suffered quantifiable "
        "damages exceeding $4.7 million as a direct result of Apex's misconduct."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # --- II. STATEMENT OF FACTS --- (5 tabs)
    heading2 = doc.add_paragraph()
    run = heading2.add_run("II.\tSTATEMENT OF UNDISPUTED FACTS")  # tab 7
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    fact_a = doc.add_paragraph()
    run = fact_a.add_run("A.\tBackground of the Parties")  # tab 8
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    fact1 = doc.add_paragraph()
    run = fact1.add_run(
        "\tMeridian is a Delaware corporation headquartered in San Jose, California, "  # tab 9
        "specializing in enterprise data analytics software. Founded in 2011, Meridian "
        "has invested over $28 million in developing its proprietary DataStream Engine, "
        "a real-time data processing framework used by more than 200 Fortune 500 companies."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    fact2 = doc.add_paragraph()
    run = fact2.add_run(
        "\tApex is a Texas limited liability company that entered the data analytics "  # tab 10
        "market in 2019. Prior to the events giving rise to this litigation, Apex had no "
        "comparable real-time processing capability and was losing significant market share "
        "to Meridian. (Declaration of Rebecca Foster, Ex. A at \u00b6\u00b6 12-15.)"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    fact_b = doc.add_paragraph()
    run = fact_b.add_run("B.\tThe Misappropriation")  # tab 11
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    fact3 = doc.add_paragraph()
    run = fact3.add_run(
        "\tIn March 2023, Dr. Nathan Whitfield, Meridian's former Lead Architect, "  # tab 12
        "resigned his position and joined Apex as Vice President of Engineering. Prior to "
        "his departure, Dr. Whitfield downloaded approximately 47,000 files from Meridian's "
        "secure repositories, including the complete source code for the DataStream Engine."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    fact4 = doc.add_paragraph()
    run = fact4.add_run(
        "\tWithin six months of Dr. Whitfield's hiring, Apex announced the launch of "  # tab 13
        "its \"RapidFlow\" platform, which independent forensic analysis confirmed shares "
        "87% code similarity with Meridian's DataStream Engine. (Expert Report of "
        "Dr. Patricia Huang, Ex. C at pp. 14-23.)"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # --- III. ARGUMENT --- (6 tabs)
    heading3 = doc.add_paragraph()
    run = heading3.add_run("III.\tARGUMENT")  # tab 14
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    arg_a = doc.add_paragraph()
    run = arg_a.add_run("A.\tLegal Standard for Summary Judgment")  # tab 15
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    arg1 = doc.add_paragraph()
    run = arg1.add_run(
        "\tSummary judgment is appropriate when \"the movant shows that there is no genuine "  # tab 16
        "dispute as to any material fact and the movant is entitled to judgment as a matter "
        "of law.\" Fed. R. Civ. P. 56(a). The Court must view the evidence in the light most "
        "favorable to the nonmoving party. Anderson v. Liberty Lobby, Inc., 477 U.S. 242, "
        "255 (1986)."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    arg_b = doc.add_paragraph()
    run = arg_b.add_run("B.\tMeridian's Algorithms Constitute Trade Secrets")  # tab 17
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    arg2 = doc.add_paragraph()
    run = arg2.add_run(
        "\tUnder the DTSA, a trade secret is information that: (1) derives independent "  # tab 18
        "economic value from not being generally known; and (2) is the subject of reasonable "
        "efforts to maintain its secrecy. 18 U.S.C. § 1839(3). Both elements are satisfied here."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    arg3 = doc.add_paragraph()
    run = arg3.add_run(
        "\tFirst, the DataStream Engine's proprietary algorithms provide Meridian with a "  # tab 19
        "significant competitive advantage, processing data at speeds 3.2 times faster than "
        "the nearest competitor. (Foster Decl. \u00b6 28.) Second, Meridian implemented "
        "comprehensive security measures, including:"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Bulleted items with tabs (4 tabs)
    bullet1 = doc.add_paragraph()
    run = bullet1.add_run("\t\tAccess restricted to authorized personnel with security clearance;")  # tabs 20, 21
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    bullet2 = doc.add_paragraph()
    run = bullet2.add_run("\t\tMulti-factor authentication and encrypted repositories;")  # tabs 22, 23
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # --- IV. CONCLUSION --- (2 tabs)
    heading4 = doc.add_paragraph()
    run = heading4.add_run("IV.\tCONCLUSION")  # tab 24
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    conclusion = doc.add_paragraph()
    run = conclusion.add_run(
        "\tFor the foregoing reasons, Plaintiff Meridian Technologies, Inc. respectfully "  # tab 25
        "requests that this Court grant its motion for summary judgment on all claims."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Signature block (no tabs)
    doc.add_paragraph()
    sig = doc.add_paragraph()
    run = sig.add_run("Respectfully submitted,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    sig2 = doc.add_paragraph()
    run = sig2.add_run("_________________________")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    sig3 = doc.add_paragraph()
    run = sig3.add_run("Katherine M. Reeves, Esq.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    sig4 = doc.add_paragraph()
    run = sig4.add_run("REEVES & CALLOWAY LLP")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    sig5 = doc.add_paragraph()
    run = sig5.add_run("Attorneys for Plaintiff")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Verify tab count
    tab_count = 0
    verify_doc = Document(OUTPUT)
    for para in verify_doc.paragraphs:
        tab_count += para.text.count('\t')
    print(f'Tab count in initial document: {tab_count}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
