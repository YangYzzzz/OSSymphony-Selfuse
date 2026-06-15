"""
Reward Script: Replace all tab characters with four spaces in a legal brief
Task ID: writer_frd_011
Domain: libreoffice_writer
Scoring:
  Precondition: Document structure preserved (paragraph count, key content markers)
  Component 1 (0.6): No tab characters remain in document
  Component 2 (0.4): Four-space sequences present as replacements (>=20, expecting ~25)
"""

import os
import re
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_frd_011'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all run texts for analysis
    all_runs_text = []
    for para in doc.paragraphs:
        for run in para.runs:
            all_runs_text.append(run.text)

    full_text = ''.join(all_runs_text)

    # Precondition gate: Document structure must be preserved
    # If the document is corrupted or fundamentally changed, score 0.0
    try:
        para_count = len(doc.paragraphs)
        normalized_text = re.sub(r'\s+', ' ', full_text).strip()
        markers = [
            "MERIDIAN TECHNOLOGIES",
            "INTRODUCTION",
            "STATEMENT OF UNDISPUTED FACTS",
            "ARGUMENT",
            "CONCLUSION",
            "24-CV-03847-RMW"
        ]
        missing = [m for m in markers if m not in normalized_text]

        if missing:
            print(f"PRECONDITION FAIL: Missing content markers: {missing}")
            print("REWARD: 0.0")
            return 0.0
        if abs(para_count - 39) > 3:
            print(f"PRECONDITION FAIL: Paragraph count {para_count} deviates too much from expected 39")
            print("REWARD: 0.0")
            return 0.0
        print(f"PRECONDITION PASS: Document structure intact ({para_count} paragraphs, all markers present)")
    except Exception as e:
        print(f"PRECONDITION ERROR: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: No tab characters remain (0.6 points)
    # Initial has 25 tabs; golden should have 0.
    try:
        tab_count = full_text.count('\t')
        if tab_count == 0:
            print(f"PASS: Component 1 - No tab characters found (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 - Found {tab_count} tab character(s), expected 0")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Four-space sequences present as replacements (0.4 points)
    # Each tab should have been replaced with 4 spaces.
    # Initial has 0 four-space sequences; golden should have ~25.
    try:
        four_space_count = 0
        for run_text in all_runs_text:
            idx = 0
            while idx <= len(run_text) - 4:
                if run_text[idx:idx + 4] == '    ':
                    four_space_count += 1
                    idx += 4
                else:
                    idx += 1

        if four_space_count >= 20:
            print(f"PASS: Component 2 - Found {four_space_count} four-space sequences (>= 20 expected) (0.4 pts)")
            total_score += 0.4
        elif four_space_count >= 10:
            partial = 0.2
            print(f"PARTIAL: Component 2 - Found {four_space_count} four-space sequences (partial credit {partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - Found {four_space_count} four-space sequences, expected >= 20")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
