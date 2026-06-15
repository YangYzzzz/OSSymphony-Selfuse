"""
Initial Setup: Insert a formatted index of tables at a specific location in the document
Task ID: writer_rd_088
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_088'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_caption_paragraph(doc, caption_text, seq_num):
    """Add a table caption paragraph with SEQ field code, mimicking LibreOffice Writer's caption format.
    Format: 'Table N: Description'
    Uses the 'Caption' style if available, otherwise normal paragraph with italic.
    """
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Apply caption styling
    run_prefix = para.add_run("Table ")
    run_prefix.italic = True
    run_prefix.font.size = Pt(10)

    # Insert SEQ field for auto-numbering: this is how LO Writer recognizes table captions
    # Field: { SEQ Table \* ARABIC }
    r_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:i/><w:sz w:val="20"/></w:rPr>'
        f'  <w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
    )
    para._element.append(r_elem)

    r_instr = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:i/><w:sz w:val="20"/></w:rPr>'
        f'  <w:instrText xml:space="preserve"> SEQ Table \\* ARABIC </w:instrText>'
        f'</w:r>'
    )
    para._element.append(r_instr)

    r_sep = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:i/><w:sz w:val="20"/></w:rPr>'
        f'  <w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
    )
    para._element.append(r_sep)

    # Cached display value
    r_val = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:i/><w:sz w:val="20"/></w:rPr>'
        f'  <w:t>{seq_num}</w:t>'
        f'</w:r>'
    )
    para._element.append(r_val)

    r_end = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:i/><w:sz w:val="20"/></w:rPr>'
        f'  <w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )
    para._element.append(r_end)

    # Add the description part
    run_desc = para.add_run(f": {caption_text}")
    run_desc.italic = True
    run_desc.font.size = Pt(10)

    return para


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ========================
    # TITLE PAGE
    # ========================
    title = doc.add_heading('Quarterly Data Analysis Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Greenfield Analytics Inc.')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Q4 2025 — Prepared by the Business Intelligence Team')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ========================
    # LIST OF TABLES SECTION (empty — agent must insert index here)
    # ========================
    doc.add_heading('List of Tables', level=1)
    # Intentionally leave this section empty — the task is to insert the index here
    doc.add_paragraph()  # empty placeholder paragraph

    # ========================
    # SECTION 1: EXECUTIVE SUMMARY
    # ========================
    doc.add_page_break()
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report provides a comprehensive analysis of Greenfield Analytics\' performance '
        'across all business units for Q4 2025. Key metrics including revenue, customer acquisition, '
        'product performance, and operational efficiency are examined in detail. The data presented '
        'spans October through December 2025 and covers all regional offices.'
    )
    doc.add_paragraph(
        'Overall, the company experienced a 12.3% increase in revenue compared to Q3 2025, '
        'driven primarily by strong performance in the Asia-Pacific region and the successful '
        'launch of the DataStream Pro product line. Customer retention rates improved to 94.7%, '
        'exceeding the annual target of 92%.'
    )

    # ========================
    # SECTION 2: SALES BY REGION
    # ========================
    doc.add_heading('2. Sales Performance by Region', level=1)
    doc.add_paragraph(
        'Regional sales data reveals significant variation across geographies. The following '
        'table summarizes total revenue, number of deals closed, and average deal size for each region.'
    )

    # TABLE 1: Sales by Region
    table1 = doc.add_table(rows=7, cols=4)
    table1.style = 'Table Grid'
    headers1 = ['Region', 'Revenue ($K)', 'Deals Closed', 'Avg Deal Size ($K)']
    data1 = [
        ['North America', '2,450', '187', '13.1'],
        ['Europe', '1,890', '142', '13.3'],
        ['Asia-Pacific', '2,180', '198', '11.0'],
        ['Latin America', '780', '64', '12.2'],
        ['Middle East & Africa', '520', '38', '13.7'],
        ['Total', '7,820', '629', '12.4'],
    ]
    for j, h in enumerate(headers1):
        cell = table1.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data1):
        for j, val in enumerate(row_data):
            table1.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Sales by Region', 1)

    doc.add_paragraph(
        'North America and Asia-Pacific together accounted for 59.2% of total revenue. '
        'The Asia-Pacific region showed the highest growth rate at 18.4% quarter-over-quarter.'
    )

    # ========================
    # SECTION 3: PRODUCT PERFORMANCE
    # ========================
    doc.add_heading('3. Product Performance', level=1)
    doc.add_paragraph(
        'Product-level analysis shows that the newly launched DataStream Pro contributed '
        'significantly to overall revenue growth. Legacy products maintained stable performance.'
    )

    # TABLE 2: Product Performance
    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    headers2 = ['Product', 'Units Sold', 'Revenue ($K)', 'Margin (%)', 'YoY Growth (%)']
    data2 = [
        ['DataStream Pro', '1,245', '3,120', '68.2', 'N/A (New)'],
        ['DataStream Basic', '2,890', '1,734', '72.5', '-4.1'],
        ['AnalyticsHub', '892', '1,784', '61.8', '8.3'],
        ['ReportForge', '1,567', '940', '55.4', '2.7'],
        ['CloudSync Module', '3,210', '642', '78.1', '15.6'],
    ]
    for j, h in enumerate(headers2):
        cell = table2.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data2):
        for j, val in enumerate(row_data):
            table2.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Product Performance', 2)

    # ========================
    # SECTION 4: CUSTOMER ACQUISITION
    # ========================
    doc.add_heading('4. Customer Acquisition and Retention', level=1)
    doc.add_paragraph(
        'Customer metrics improved significantly this quarter. New customer acquisition '
        'increased while churn rates decreased, resulting in a net positive growth trajectory.'
    )

    # TABLE 3: Customer Acquisition by Channel
    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'
    headers3 = ['Channel', 'New Customers', 'CAC ($)', 'Conversion Rate (%)']
    data3 = [
        ['Organic Search', '342', '45', '3.8'],
        ['Paid Advertising', '518', '128', '2.1'],
        ['Partner Referrals', '267', '62', '5.4'],
        ['Direct Sales', '189', '210', '8.7'],
        ['Social Media', '156', '85', '1.9'],
    ]
    for j, h in enumerate(headers3):
        cell = table3.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data3):
        for j, val in enumerate(row_data):
            table3.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Customer Acquisition by Channel', 3)

    doc.add_paragraph(
        'Partner referrals showed the highest conversion rate at 5.4%, while paid advertising '
        'generated the highest volume of new customers. The blended customer acquisition cost '
        'was $94 per customer, down from $108 in Q3.'
    )

    # TABLE 4: Customer Retention Metrics
    table4 = doc.add_table(rows=5, cols=4)
    table4.style = 'Table Grid'
    headers4 = ['Segment', 'Retention Rate (%)', 'Churn Rate (%)', 'NPS Score']
    data4 = [
        ['Enterprise', '97.2', '2.8', '72'],
        ['Mid-Market', '94.1', '5.9', '65'],
        ['SMB', '89.3', '10.7', '58'],
        ['Startup', '82.6', '17.4', '61'],
    ]
    for j, h in enumerate(headers4):
        cell = table4.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data4):
        for j, val in enumerate(row_data):
            table4.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Customer Retention Metrics', 4)

    # ========================
    # SECTION 5: FINANCIAL OVERVIEW
    # ========================
    doc.add_page_break()
    doc.add_heading('5. Financial Overview', level=1)
    doc.add_paragraph(
        'The financial performance of Greenfield Analytics in Q4 2025 exceeded expectations. '
        'Operating margins expanded due to improved operational efficiency and favorable '
        'product mix shifts toward higher-margin offerings.'
    )

    # TABLE 5: Quarterly Financial Summary
    table5 = doc.add_table(rows=8, cols=3)
    table5.style = 'Table Grid'
    headers5 = ['Metric', 'Q4 2025', 'Q3 2025']
    data5 = [
        ['Total Revenue', '$7,820K', '$6,964K'],
        ['Cost of Goods Sold', '$2,581K', '$2,369K'],
        ['Gross Profit', '$5,239K', '$4,595K'],
        ['Operating Expenses', '$3,450K', '$3,280K'],
        ['Operating Income', '$1,789K', '$1,315K'],
        ['Net Income', '$1,432K', '$1,052K'],
        ['Earnings Per Share', '$2.14', '$1.57'],
    ]
    for j, h in enumerate(headers5):
        cell = table5.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data5):
        for j, val in enumerate(row_data):
            table5.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Quarterly Financial Summary', 5)

    # TABLE 6: Expense Breakdown
    table6 = doc.add_table(rows=7, cols=3)
    table6.style = 'Table Grid'
    headers6 = ['Category', 'Amount ($K)', 'Percentage of Revenue (%)']
    data6 = [
        ['R&D', '1,250', '16.0'],
        ['Sales & Marketing', '1,120', '14.3'],
        ['General & Administrative', '580', '7.4'],
        ['Customer Support', '320', '4.1'],
        ['Infrastructure', '180', '2.3'],
        ['Total OpEx', '3,450', '44.1'],
    ]
    for j, h in enumerate(headers6):
        cell = table6.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data6):
        for j, val in enumerate(row_data):
            table6.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Expense Breakdown', 6)

    doc.add_paragraph(
        'R&D expenditure increased by 8.2% to support the DataStream Pro launch and upcoming '
        'features for AnalyticsHub. Sales & Marketing spend remained stable as a percentage '
        'of revenue, indicating improved efficiency in customer acquisition.'
    )

    # ========================
    # SECTION 6: OPERATIONAL METRICS
    # ========================
    doc.add_heading('6. Operational Metrics', level=1)
    doc.add_paragraph(
        'Operational performance indicators demonstrate continued improvement in service '
        'delivery and platform reliability. The engineering team achieved significant '
        'reductions in incident response times.'
    )

    # TABLE 7: Platform Uptime and Performance
    table7 = doc.add_table(rows=5, cols=4)
    table7.style = 'Table Grid'
    headers7 = ['Service', 'Uptime (%)', 'Avg Response (ms)', 'P99 Latency (ms)']
    data7 = [
        ['DataStream API', '99.97', '42', '185'],
        ['AnalyticsHub Dashboard', '99.92', '128', '420'],
        ['ReportForge Engine', '99.89', '315', '890'],
        ['CloudSync', '99.99', '18', '65'],
    ]
    for j, h in enumerate(headers7):
        cell = table7.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data7):
        for j, val in enumerate(row_data):
            table7.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Platform Uptime and Performance', 7)

    doc.add_paragraph(
        'CloudSync maintained the highest uptime at 99.99%, reflecting the team\'s investment '
        'in redundancy and automated failover systems during Q3.'
    )

    # TABLE 8: Support Ticket Summary
    table8 = doc.add_table(rows=6, cols=4)
    table8.style = 'Table Grid'
    headers8 = ['Priority', 'Tickets Opened', 'Tickets Resolved', 'Avg Resolution (hrs)']
    data8 = [
        ['Critical (P1)', '12', '12', '2.4'],
        ['High (P2)', '87', '84', '8.1'],
        ['Medium (P3)', '342', '328', '24.6'],
        ['Low (P4)', '518', '497', '48.2'],
        ['Total', '959', '921', '28.7'],
    ]
    for j, h in enumerate(headers8):
        cell = table8.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row_data in enumerate(data8):
        for j, val in enumerate(row_data):
            table8.cell(i + 1, j).text = val
    add_caption_paragraph(doc, 'Support Ticket Summary', 8)

    # ========================
    # SECTION 7: CONCLUSION
    # ========================
    doc.add_heading('7. Conclusion and Outlook', level=1)
    doc.add_paragraph(
        'Q4 2025 marked a strong finish to the fiscal year for Greenfield Analytics. The '
        'successful launch of DataStream Pro, combined with improved retention rates and '
        'operational efficiency, positions the company well for continued growth in 2026. '
        'Key priorities for Q1 2026 include expanding the partner ecosystem, launching the '
        'AnalyticsHub 3.0 upgrade, and entering the Japanese market.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
