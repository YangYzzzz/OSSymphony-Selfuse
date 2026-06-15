"""
Initial Setup: Set page border padding to 2cm from page edge
Task ID: writer_page_048
Domain: libreoffice_writer

Creates framed_announcement.docx with:
- 1-page announcement content
- A4 portrait, margins 2.54cm all sides
- Page border: solid blue (#0000FF), 1pt, offsetFrom='page'
- Initial padding: 0.5cm from page edge
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_page_048'
OUTPUT = f'{WORKDIR}/framed_announcement.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_borders(section, space_pt, color_hex='0000FF', sz=8, offset_from='page'):
    """
    Add page borders to a section.

    Args:
        section: docx section object
        space_pt: distance in points (integer) from page edge to border
        color_hex: border color as hex string (no #)
        sz: border width in 1/8 pt units (8 = 1pt)
        offset_from: 'page' or 'text'
    """
    sectPr = section._sectPr

    # Remove existing pgBorders if any
    existing = sectPr.find(qn('w:pgBorders'))
    if existing is not None:
        sectPr.remove(existing)

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), offset_from)

    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), str(space_pt))
        border.set(qn('w:color'), color_hex)
        pgBorders.append(border)

    # Insert pgBorders before pgSz or at end
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        sectPr.insert(list(sectPr).index(pgSz), pgBorders)
    else:
        sectPr.append(pgBorders)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()
    section = doc.sections[0]

    # --- Page setup: A4 portrait ---
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Page border: solid blue 1pt, padding 0.5cm from page edge ---
    # 0.5cm = 14.17pt → 14 pt
    space_initial = 14  # ~0.5cm in points
    add_page_borders(section, space_pt=space_initial, color_hex='0000FF', sz=8, offset_from='page')

    # --- Announcement content ---
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('ANNUAL COMMUNITY AWARDS CEREMONY')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)

    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Celebrating Excellence and Achievement')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    # Event details
    details = [
        ('Date:', 'Saturday, 15th November 2025'),
        ('Time:', '7:00 PM – 10:00 PM'),
        ('Venue:', 'Grand Ballroom, Riverside Convention Centre'),
        ('Dress Code:', 'Black Tie / Formal Attire'),
    ]

    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r_label = p.add_run(label + ' ')
        r_label.bold = True
        r_label.font.size = Pt(11)
        r_value = p.add_run(value)
        r_value.font.size = Pt(11)

    doc.add_paragraph()

    # Body text
    body_text = (
        'Join us for an extraordinary evening as we honour the outstanding contributions '
        'of our community members across the fields of education, arts, science, '
        'public service, and entrepreneurship. This year\'s ceremony will feature '
        'live performances, a gala dinner, and the presentation of twelve prestigious awards.'
    )
    body = doc.add_paragraph(body_text)
    body.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in body.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()

    # RSVP section
    rsvp = doc.add_paragraph()
    rsvp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = rsvp.add_run('RSVP by 31st October 2025')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = contact.add_run('events@riversidemuni.gov.au  |  (03) 9854 2200')
    r.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
