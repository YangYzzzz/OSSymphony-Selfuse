"""
Initial Setup: Train schedule document with no tabstops
Task ID: osworld_writer_tabstop_007
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_tabstop_007'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Train schedule data — 30 realistic lines
    # Format: "TRAIN_ID HH:MM Route Platform N"
    # No tabstops, no tab characters — plain left-aligned text
    train_data = [
        'TR001 07:30 London-Birmingham Platform 3',
        'TR002 07:45 London-Bristol Platform 1',
        'TR003 08:00 London-Liverpool Platform 6',
        'TR004 08:15 London-Manchester Platform 7',
        'TR005 08:30 London-Leeds Platform 4',
        'TR006 08:45 London-Newcastle Platform 2',
        'TR007 09:00 London-Edinburgh Platform 9',
        'TR008 09:15 London-Glasgow Platform 5',
        'TR009 09:30 London-Cardiff Platform 8',
        'TR010 09:45 London-Sheffield Platform 3',
        'TR011 10:00 London-Nottingham Platform 1',
        'TR012 10:15 London-Leicester Platform 6',
        'TR013 10:30 London-Derby Platform 4',
        'TR014 10:45 London-Coventry Platform 2',
        'TR015 11:00 London-Oxford Platform 7',
        'TR016 11:15 London-Cambridge Platform 5',
        'TR017 11:30 London-Norwich Platform 9',
        'TR018 11:45 London-Brighton Platform 3',
        'TR019 12:00 London-Southampton Platform 1',
        'TR020 12:15 London-Portsmouth Platform 6',
        'TR021 12:30 London-Exeter Platform 4',
        'TR022 12:45 London-Plymouth Platform 2',
        'TR023 13:00 London-Bath Platform 8',
        'TR024 13:15 London-Reading Platform 5',
        'TR025 13:30 London-Swindon Platform 3',
        'TR026 13:45 London-Peterborough Platform 7',
        'TR027 14:00 London-York Platform 9',
        'TR028 14:15 London-Hull Platform 1',
        'TR029 14:30 London-Bradford Platform 4',
        'TR030 14:45 London-Stoke-on-Trent Platform 6',
    ]

    for line in train_data:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(11)
        # No tabstops added, no tab characters inserted — plain left-aligned

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
