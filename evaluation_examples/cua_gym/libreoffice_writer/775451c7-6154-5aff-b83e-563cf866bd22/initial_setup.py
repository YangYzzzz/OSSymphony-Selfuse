"""
Initial Setup: Image adjustment instructions in notes.docx, illustration.png on Desktop
Task ID: osworld_multi_apps_writer_gimp_066
Domain: libreoffice_writer + gimp (multi-app)
"""

import os
import shlex
import subprocess
import time
import numpy as np

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_writer_gimp_066'
DESKTOP = '/home/user/Desktop'
DOC_OUTPUT = f'{WORKDIR}/notes.docx'
IMG_OUTPUT = f'{DESKTOP}/illustration.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_notes_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading('Image Processing Instructions', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction paragraph
    doc.add_paragraph(
        'Follow the steps below carefully to apply the posterize effect and color reduction '
        'to the illustration image. Save the final result as a new file.'
    )

    doc.add_paragraph('')

    # Step-by-step instructions heading
    doc.add_heading('Steps to Apply Poster Art Effect', level=2)

    # Step 1
    step1 = doc.add_paragraph(style='List Number')
    run1 = step1.add_run('Open illustration.png')
    run1.bold = True
    step1.add_run(' from the Desktop in GIMP.')

    # Step 2
    step2 = doc.add_paragraph(style='List Number')
    run2 = step2.add_run('Apply Posterize Effect:')
    run2.bold = True
    step2.add_run(
        ' In the menu, go to Colors > Posterize. Set the number of posterize levels to '
        '4 and click OK.'
    )

    # Step 3
    step3 = doc.add_paragraph(style='List Number')
    run3 = step3.add_run('Convert to Indexed Mode:')
    run3.bold = True
    step3.add_run(
        ' Go to Image > Mode > Indexed. In the dialog, set the maximum number of colors '
        'to 32. Under the "Generate optimum palette" option, select '
        '"Make Palette Colorful". Click Convert.'
    )

    # Step 4
    step4 = doc.add_paragraph(style='List Number')
    run4 = step4.add_run('Convert back to RGB Mode:')
    run4.bold = True
    step4.add_run(
        ' Go to Image > Mode > RGB to convert the indexed image back to RGB color mode.'
    )

    # Step 5
    step5 = doc.add_paragraph(style='List Number')
    run5 = step5.add_run('Export as PNG:')
    run5.bold = True
    step5.add_run(
        ' Go to File > Export As. Save the file as '
        '"illustration_poster.png" on the Desktop. Click Export and then '
        '"Export" again to confirm PNG settings.'
    )

    doc.add_paragraph('')

    # Note section
    doc.add_heading('Important Notes', level=2)
    note = doc.add_paragraph()
    note_run = note.add_run('Note: ')
    note_run.bold = True
    note.add_run(
        'Do NOT overwrite the original illustration.png. Save the result as '
        'a separate file named illustration_poster.png on the Desktop.'
    )

    # Summary table
    doc.add_paragraph('')
    doc.add_heading('Parameter Summary', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    headers = [('Parameter', 'Value'),
               ('Posterize Levels', '4'),
               ('Color Mode', 'Indexed → RGB'),
               ('Number of Colors', '32')]

    for i, (param, value) in enumerate(headers):
        row = table.rows[i]
        row.cells[0].text = param
        row.cells[1].text = value
        if i == 0:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.save(DOC_OUTPUT)
    print(f'Notes document created: {DOC_OUTPUT}')


def create_illustration_png():
    from PIL import Image, ImageDraw
    import numpy as np

    # Create a colorful illustration image (800x600)
    width, height = 800, 600
    img = Image.new('RGB', (width, height), (240, 245, 250))
    draw = ImageDraw.Draw(img)

    # Sky gradient background (simulated)
    for y in range(height // 2):
        r = int(135 + (y / (height // 2)) * 50)
        g = int(180 + (y / (height // 2)) * 30)
        b = int(230 - (y / (height // 2)) * 60)
        draw.line([(0, y), (width, y)], fill=(r, g, b))

    # Ground / grass
    for y in range(height // 2, height):
        factor = (y - height // 2) / (height // 2)
        r = int(60 + factor * 20)
        g = int(130 + factor * 30)
        b = int(40 + factor * 15)
        draw.line([(0, y), (width, y)], fill=(r, g, b))

    # Sun
    draw.ellipse([600, 50, 720, 170], fill=(255, 220, 50), outline=(255, 180, 0), width=3)

    # Mountains
    mountain_coords = [
        [(100, 350), (220, 160), (340, 350)],
        [(260, 350), (400, 110), (540, 350)],
        [(460, 350), (580, 175), (700, 350)],
    ]
    mountain_colors = [
        (120, 100, 160),
        (90, 80, 140),
        (110, 95, 150),
    ]
    for coords, color in zip(mountain_coords, mountain_colors):
        draw.polygon(coords, fill=color, outline=(80, 60, 100))

    # Snow caps
    snow_caps = [
        [(190, 185), (220, 160), (250, 185)],
        [(370, 135), (400, 110), (430, 135)],
        [(555, 200), (580, 175), (605, 200)],
    ]
    for cap in snow_caps:
        draw.polygon(cap, fill=(255, 255, 255))

    # Trees (triangles)
    tree_positions = [70, 140, 560, 630, 700, 740]
    for tx in tree_positions:
        ty = 340
        tree_h = 80
        draw.polygon([(tx, ty), (tx - 25, ty + tree_h), (tx + 25, ty + tree_h)],
                     fill=(30, 100, 50), outline=(20, 70, 30))
        draw.polygon([(tx, ty + 35), (tx - 30, ty + tree_h + 30), (tx + 30, ty + tree_h + 30)],
                     fill=(25, 85, 40), outline=(15, 60, 25))
        draw.rectangle([tx - 8, ty + tree_h + 30, tx + 8, ty + tree_h + 60],
                       fill=(100, 70, 40))

    # River
    river_pts = [(200, height), (230, 500), (260, 450), (300, 420),
                 (350, 410), (400, 415), (440, 430), (480, 460), (510, height)]
    draw.polygon(river_pts, fill=(60, 140, 200))

    # Flowers
    flower_data = [(150, 490, (255, 80, 80)), (350, 510, (255, 200, 50)),
                   (500, 480, (200, 100, 230)), (660, 500, (255, 120, 50)),
                   (720, 460, (100, 200, 255))]
    for fx, fy, fc in flower_data:
        draw.ellipse([fx - 12, fy - 12, fx + 12, fy + 12], fill=fc)
        draw.ellipse([fx - 5, fy - 5, fx + 5, fy + 5], fill=(255, 255, 100))

    # Clouds
    cloud_positions = [(150, 80), (380, 60), (520, 90)]
    for cx, cy in cloud_positions:
        for dx, dy, r in [(0, 0, 28), (25, -10, 22), (-25, -8, 20), (15, 12, 18), (-18, 10, 16)]:
            draw.ellipse([cx + dx - r, cy + dy - r, cx + dx + r, cy + dy + r],
                         fill=(255, 255, 255))

    os.makedirs(DESKTOP, exist_ok=True)
    img.save(IMG_OUTPUT, 'PNG')
    print(f'Illustration image created: {IMG_OUTPUT}')


def remove_golden_if_exists():
    """Remove any pre-existing poster file to ensure clean initial state."""
    poster_path = f'{DESKTOP}/illustration_poster.png'
    if os.path.exists(poster_path):
        os.remove(poster_path)
        print(f'Removed pre-existing file: {poster_path}')


def main():
    create_notes_docx()
    create_illustration_png()
    remove_golden_if_exists()

    # GUI-ready startup: open notes.docx in LibreOffice Writer first
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=3.0)
    # Also open illustration.png in GIMP so it's ready for editing
    launch_gui(f'gimp "{IMG_OUTPUT}"', delay_sec=2.0)

    print('GUI_READY: launched LibreOffice Writer (notes.docx) and GIMP (illustration.png) with DISPLAY=:0')


main()
