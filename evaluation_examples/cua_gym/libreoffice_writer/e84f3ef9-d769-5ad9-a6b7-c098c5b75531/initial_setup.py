"""
Initial Setup: Create a Writer document with a System Requirements heading
and plain text listing requirements (no table).
Task ID: writer_tech_019
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_019'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('CloudSync Pro - Technical Documentation', level=0)

    # Introduction paragraph
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'CloudSync Pro is an enterprise-grade file synchronization platform '
        'designed for organizations with distributed teams. It provides real-time '
        'file syncing across multiple devices while maintaining end-to-end encryption '
        'and compliance with industry standards including SOC 2 Type II and GDPR.'
    )

    # System Requirements section - plain text only, NO table
    doc.add_heading('System Requirements', level=1)
    doc.add_paragraph(
        'Before deploying CloudSync Pro, ensure that your infrastructure meets '
        'the following minimum and recommended specifications. The application has '
        'been tested across a variety of hardware configurations to ensure broad '
        'compatibility.'
    )
    doc.add_paragraph(
        'Processor: The application requires at minimum a dual-core processor '
        'clocked at 2.0 GHz or higher. For optimal performance, especially in '
        'environments handling large file batches, a quad-core processor at 3.0 GHz '
        'or above is recommended.'
    )
    doc.add_paragraph(
        'Memory: A minimum of 4 GB RAM is required for basic operation. '
        'Organizations expecting concurrent usage by more than 50 users should '
        'provision at least 16 GB RAM to maintain responsive performance.'
    )
    doc.add_paragraph(
        'Storage: At least 20 GB of available disk space is needed for the '
        'application and its local cache. Production deployments should allocate '
        '100 GB or more depending on the expected data volume.'
    )

    # Additional sections for realism
    doc.add_heading('Installation Guide', level=1)
    doc.add_paragraph(
        'Download the latest installer package from the CloudSync Pro admin portal. '
        'Run the installer with administrator privileges and follow the on-screen '
        'wizard. The installation typically completes within 5-10 minutes depending '
        'on system speed.'
    )
    doc.add_paragraph(
        'After installation, launch the configuration utility to connect to your '
        'organization\'s CloudSync Pro server. You will need the server URL and an '
        'admin API key, both of which can be obtained from the IT administrator.'
    )

    doc.add_heading('Network Configuration', level=1)
    doc.add_paragraph(
        'CloudSync Pro communicates over HTTPS (port 443) by default. Ensure that '
        'your firewall allows outbound connections to the sync endpoint. For '
        'on-premises deployments, ports 8443 and 8444 must also be open for the '
        'admin console and WebSocket connections respectively.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
