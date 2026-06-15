"""
Reward Script: Add a 3-column, 2-row system requirements table
Task ID: writer_tech_019
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table exists with 3 columns
  Component 2 (0.3): Table has exactly 2 rows (header + 1 data row)
  Component 3 (0.2): Header row contains Component, Minimum, Recommended
  Component 4 (0.2): Data row has non-empty content in all 3 cells
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_019'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: file must have tables
    if len(doc.tables) == 0:
        print("FAIL: No tables found in the document")
        print("REWARD: 0.0")
        return 0.0

    # Find the target table — look for a table with 3 columns
    target_table = None
    for t in doc.tables:
        if len(t.columns) == 3:
            target_table = t
            break

    # Component 1: Table exists with 3 columns (0.3 points)
    try:
        if target_table is not None:
            print(f"PASS: Component 1 — Found table with 3 columns (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No table with exactly 3 columns found. Tables have columns: {[len(t.columns) for t in doc.tables]}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    if target_table is None:
        # Can't check further components without the right table
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Table has exactly 2 rows (header + 1 data row) (0.3 points)
    try:
        num_rows = len(target_table.rows)
        if num_rows == 2:
            print(f"PASS: Component 2 — Table has exactly 2 rows (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Expected 2 rows, found {num_rows}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Header row contains Component, Minimum, Recommended (0.2 points)
    try:
        header_cells = [cell.text.strip().lower() for cell in target_table.rows[0].cells]
        expected_headers = ["component", "minimum", "recommended"]
        if header_cells == expected_headers:
            print(f"PASS: Component 3 — Header row matches: {header_cells} (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Expected headers {expected_headers}, found {header_cells}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Data row (row 1) has non-empty content in all 3 cells (0.2 points)
    try:
        if len(target_table.rows) >= 2:
            data_cells = [cell.text.strip() for cell in target_table.rows[1].cells]
            all_filled = all(len(c) > 0 for c in data_cells)
            if all_filled:
                print(f"PASS: Component 4 — Data row has content in all cells: {data_cells} (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 4 — Some data cells are empty: {data_cells}")
        else:
            print(f"FAIL: Component 4 — No data row exists (only {len(target_table.rows)} row(s))")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
