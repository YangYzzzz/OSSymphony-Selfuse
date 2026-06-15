"""
Initial Setup: English-only employee safety manual
Task ID: writer_hr_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_072'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body_paragraph(doc, text, bold=False):
    """Add a body paragraph with optional bold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    return p


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_heading('Employee Safety Manual', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Workplace Health & Safety Department')
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Effective Date: January 15, 2025')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    doc.add_paragraph()  # spacing

    # --- Table of Contents placeholder ---
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Personal Protective Equipment (PPE)',
        '2. Fire Safety and Prevention',
        '3. Chemical Handling Procedures',
        '4. Electrical Safety',
        '5. Fall Protection',
        '6. Confined Space Entry',
        '7. Lockout/Tagout (LOTO) Procedures',
        '8. Forklift and Heavy Equipment Safety',
        '9. Emergency First Aid Response',
        '10. Workplace Ergonomics',
        'Emergency Contacts',
        'Safety Equipment Checklist',
        'Emergency Evacuation Maps',
    ]
    for item in toc_items:
        add_body_paragraph(doc, item)

    doc.add_page_break()

    # --- Procedure 1: PPE ---
    add_heading_styled(doc, '1. Personal Protective Equipment (PPE)', level=1)
    add_body_paragraph(doc, 'All employees must wear appropriate personal protective equipment when entering designated hazard zones.')
    steps_ppe = [
        'Step 1: Identify the hazard type in the work area by consulting the posted safety signage.',
        'Step 2: Select the appropriate PPE from the equipment station, including hard hats, safety goggles, gloves, and steel-toed boots.',
        'Step 3: Inspect all PPE for visible damage, cracks, or wear before use.',
        'Step 4: Ensure a proper fit for each piece of equipment. Adjust straps and closures as needed.',
        'Step 5: Report any damaged or defective PPE to your supervisor immediately and obtain a replacement.',
    ]
    for step in steps_ppe:
        add_body_paragraph(doc, step)

    # --- Procedure 2: Fire Safety ---
    add_heading_styled(doc, '2. Fire Safety and Prevention', level=1)
    add_body_paragraph(doc, 'Understanding fire hazards and prevention techniques is critical for all personnel.')
    steps_fire = [
        'Step 1: Familiarize yourself with the location of all fire extinguishers, fire alarms, and emergency exits on your floor.',
        'Step 2: Keep all workspaces free of flammable materials, oily rags, and combustible waste.',
        'Step 3: If you discover a fire, activate the nearest fire alarm pull station immediately.',
        'Step 4: Attempt to extinguish small fires only if you have been trained and the fire is contained. Use the PASS technique: Pull, Aim, Squeeze, Sweep.',
        'Step 5: Evacuate via the nearest emergency exit if the fire cannot be controlled. Do not use elevators.',
    ]
    for step in steps_fire:
        add_body_paragraph(doc, step)

    # --- Procedure 3: Chemical Handling ---
    add_heading_styled(doc, '3. Chemical Handling Procedures', level=1)
    add_body_paragraph(doc, 'Proper chemical handling minimizes exposure risks and prevents environmental contamination.')
    steps_chem = [
        'Step 1: Review the Safety Data Sheet (SDS) for each chemical before handling.',
        'Step 2: Wear appropriate PPE including chemical-resistant gloves, splash goggles, and a lab coat.',
        'Step 3: Use chemicals only in well-ventilated areas or under fume hoods.',
        'Step 4: Never mix chemicals unless specifically instructed in the approved procedure.',
        'Step 5: Dispose of chemical waste in designated containers following hazardous waste protocols.',
    ]
    for step in steps_chem:
        add_body_paragraph(doc, step)

    # --- Procedure 4: Electrical Safety ---
    add_heading_styled(doc, '4. Electrical Safety', level=1)
    add_body_paragraph(doc, 'Electrical hazards can cause severe injuries including burns, shock, and electrocution.')
    steps_elec = [
        'Step 1: Only qualified electricians may work on electrical systems rated above 50 volts.',
        'Step 2: Inspect power cords, plugs, and outlets for damage before use.',
        'Step 3: Never overload electrical outlets or use extension cords as permanent wiring.',
        'Step 4: De-energize equipment before performing maintenance using lockout/tagout procedures.',
        'Step 5: Report any electrical hazards such as exposed wiring, sparking outlets, or burning smells to maintenance immediately.',
    ]
    for step in steps_elec:
        add_body_paragraph(doc, step)

    # --- Procedure 5: Fall Protection ---
    add_heading_styled(doc, '5. Fall Protection', level=1)
    add_body_paragraph(doc, 'Falls are the leading cause of workplace fatalities in construction and industrial settings.')
    steps_fall = [
        'Step 1: Use guardrails, safety nets, or personal fall arrest systems when working at heights above 6 feet.',
        'Step 2: Inspect harnesses, lanyards, and anchor points before each use.',
        'Step 3: Maintain three points of contact when climbing ladders.',
        'Step 4: Keep work areas free of tripping hazards including loose cables, tools, and debris.',
        'Step 5: Report any damaged fall protection equipment and remove it from service immediately.',
    ]
    for step in steps_fall:
        add_body_paragraph(doc, step)

    doc.add_page_break()

    # --- Procedure 6: Confined Space Entry ---
    add_heading_styled(doc, '6. Confined Space Entry', level=1)
    add_body_paragraph(doc, 'Confined spaces present unique atmospheric and physical hazards requiring specialized protocols.')
    steps_confined = [
        'Step 1: Obtain a confined space entry permit from the safety coordinator before entry.',
        'Step 2: Test the atmosphere for oxygen levels, combustible gases, and toxic substances using a calibrated monitor.',
        'Step 3: Ensure continuous ventilation is in place and functioning properly.',
        'Step 4: Station an attendant outside the confined space who maintains visual or voice contact at all times.',
        'Step 5: Have rescue equipment and trained rescue personnel available on-site before entry begins.',
    ]
    for step in steps_confined:
        add_body_paragraph(doc, step)

    # --- Procedure 7: LOTO ---
    add_heading_styled(doc, '7. Lockout/Tagout (LOTO) Procedures', level=1)
    add_body_paragraph(doc, 'Lockout/tagout procedures prevent unexpected startup of machinery during maintenance.')
    steps_loto = [
        'Step 1: Notify all affected employees that a lockout/tagout procedure is about to begin.',
        'Step 2: Shut down the equipment using the normal stopping procedure.',
        'Step 3: Isolate all energy sources including electrical, hydraulic, pneumatic, and thermal.',
        'Step 4: Apply your personal lock and tag to each energy isolation device.',
        'Step 5: Verify zero energy state by attempting to restart the equipment before beginning work.',
    ]
    for step in steps_loto:
        add_body_paragraph(doc, step)

    # --- Procedure 8: Forklift Safety ---
    add_heading_styled(doc, '8. Forklift and Heavy Equipment Safety', level=1)
    add_body_paragraph(doc, 'Only trained and certified operators may operate forklifts and heavy equipment.')
    steps_forklift = [
        'Step 1: Complete a pre-operation inspection checking brakes, steering, horn, lights, and hydraulic systems.',
        'Step 2: Wear a seatbelt at all times while operating the forklift.',
        'Step 3: Observe all posted speed limits and travel with forks lowered to ground level.',
        'Step 4: Sound the horn at intersections, blind corners, and when approaching pedestrians.',
        'Step 5: Never exceed the rated load capacity. Check the load chart before lifting.',
    ]
    for step in steps_forklift:
        add_body_paragraph(doc, step)

    # --- Procedure 9: First Aid ---
    add_heading_styled(doc, '9. Emergency First Aid Response', level=1)
    add_body_paragraph(doc, 'Quick and proper first aid response can save lives and reduce injury severity.')
    steps_firstaid = [
        'Step 1: Assess the scene for safety before approaching the injured person.',
        'Step 2: Call emergency services (911) for life-threatening injuries or conditions.',
        'Step 3: Apply direct pressure to control bleeding using clean cloth or sterile dressings.',
        'Step 4: For suspected spinal injuries, do not move the victim. Stabilize the head and neck.',
        'Step 5: Document the incident on the Workplace Injury Report form within 24 hours.',
    ]
    for step in steps_firstaid:
        add_body_paragraph(doc, step)

    # --- Procedure 10: Ergonomics ---
    add_heading_styled(doc, '10. Workplace Ergonomics', level=1)
    add_body_paragraph(doc, 'Proper ergonomics prevent repetitive strain injuries and musculoskeletal disorders.')
    steps_ergo = [
        'Step 1: Adjust your chair height so your feet rest flat on the floor and thighs are parallel to the ground.',
        'Step 2: Position your monitor at arm\'s length with the top of the screen at eye level.',
        'Step 3: Keep your wrists in a neutral position while typing. Use a wrist rest if needed.',
        'Step 4: Take a 5-minute break every hour to stretch and change positions.',
        'Step 5: Report any discomfort, numbness, or pain to your supervisor and request an ergonomic assessment.',
    ]
    for step in steps_ergo:
        add_body_paragraph(doc, step)

    doc.add_page_break()

    # --- Emergency Contacts ---
    add_heading_styled(doc, 'Emergency Contacts', level=1)
    add_body_paragraph(doc, 'Contact the following personnel and services in case of an emergency:')
    contacts = [
        'Emergency Services: 911',
        'Plant Safety Manager - Robert Alvarez: (555) 234-5678',
        'Health & Safety Coordinator - Linda Park: (555) 345-6789',
        'Environmental Compliance - James Whitfield: (555) 456-7890',
        'Fire Department (Non-Emergency) - Station 42: (555) 567-8901',
        'Poison Control Center: (800) 222-1222',
        'Corporate Security - Night Shift: (555) 678-9012',
        'Human Resources - Maria Santos: (555) 789-0123',
    ]
    for c in contacts:
        add_body_paragraph(doc, c)

    # --- Safety Equipment Checklist ---
    add_heading_styled(doc, 'Safety Equipment Checklist', level=1)
    add_body_paragraph(doc, 'The following equipment must be inspected and maintained according to schedule:')
    equipment = [
        'Hard Hats - Inspect monthly, replace every 5 years',
        'Safety Goggles - Inspect before each use, replace if scratched or damaged',
        'Chemical-Resistant Gloves - Inspect before each use, replace if torn or degraded',
        'Steel-Toed Boots - Inspect monthly, replace when sole tread is worn',
        'Fire Extinguishers (ABC Type) - Inspect monthly, professional service annually',
        'First Aid Kits - Inventory quarterly, restock as needed',
        'Spill Containment Kits - Inspect quarterly, replace used absorbents',
        'Fall Protection Harnesses - Inspect before each use, annual professional inspection',
        'Respirators (N95 and Half-Face) - Inspect before each use, replace cartridges per schedule',
        'AED Defibrillators - Test monthly, replace pads every 2 years',
    ]
    for eq in equipment:
        add_body_paragraph(doc, eq)

    # --- Evacuation Maps ---
    add_heading_styled(doc, 'Emergency Evacuation Maps', level=1)
    add_body_paragraph(doc, 'Evacuation route maps are posted at all exits and common areas. Review the maps for your designated floor and assembly point.')
    add_body_paragraph(doc, 'Assembly Point A: North Parking Lot (Building 1, Floors 1-3)')
    add_body_paragraph(doc, 'Assembly Point B: South Garden Area (Building 1, Floors 4-6)')
    add_body_paragraph(doc, 'Assembly Point C: East Visitor Lot (Building 2, All Floors)')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
