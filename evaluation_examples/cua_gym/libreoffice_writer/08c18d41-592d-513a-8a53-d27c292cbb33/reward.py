"""
Reward Script: Bilingual (English/Spanish) Employee Safety Manual
Task ID: writer_hr_072
Domain: libreoffice_writer

Scoring Rubric:
  Component 1: Two-column table layout exists (0.25 pts)
      - Golden has 28 tables, initial has 0
  Component 2: 10 safety procedures with bilingual steps in tables (0.30 pts)
      - Each procedure has a steps table with 6 rows (header + 5 steps), 2 cols
      - Spanish steps start with "Paso"
  Component 3: Spanish text is italic, English is not (0.20 pts)
      - Right column (col 1) runs have italic=True
      - Left column (col 0) runs do NOT have italic=True
  Component 4: Emergency contacts table with bilingual headers (0.10 pts)
      - 3-column table with contact/name/phone, 8+ data rows
  Component 5: Safety equipment checklist table (0.10 pts)
      - 2-column table with equipment and schedule, 10+ rows
  Component 6: Emergency evacuation map placeholder (0.05 pts)
      - A table/cell containing "EVACUATION MAP" or similar placeholder text
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_072'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    tables = doc.tables

    # Component 1: Two-column table layout exists (0.25 pts)
    # Initial has 0 tables, golden has 28. We require at least 20 tables
    # (10 procedures * 2 tables each = 20, plus contacts, checklist, etc.)
    try:
        num_tables = len(tables)
        if num_tables >= 20:
            print(f"PASS: Component 1 — Document has {num_tables} tables (>= 20 required) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Document has {num_tables} tables, expected >= 20")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: 10 safety procedures with bilingual steps in tables (0.30 pts)
    # Each procedure should have a steps table: 6 rows (1 header + 5 steps), 2 columns
    # The header row should say "English" / "Espanol" and steps should have "Step"/"Paso"
    try:
        procedure_step_tables = 0
        for t in tables:
            try:
                num_rows = len(t.rows)
                num_cols = len(t.columns)
                if num_rows == 6 and num_cols == 2:
                    # Check header row
                    header_left = t.cell(0, 0).text.strip().lower()
                    header_right = t.cell(0, 1).text.strip().lower()
                    if 'english' in header_left and 'espa' in header_right:
                        # Check that step rows have "Step" in left and "Paso" in right
                        has_steps = False
                        left_text = t.cell(1, 0).text.strip()
                        right_text = t.cell(1, 1).text.strip()
                        if left_text.startswith('Step') and right_text.startswith('Paso'):
                            has_steps = True
                        if has_steps:
                            procedure_step_tables += 1
            except Exception:
                continue

        # Award partial credit: 0.03 per procedure step table found, max 0.30
        step_score = min(procedure_step_tables * 0.03, 0.30)
        if procedure_step_tables >= 10:
            print(f"PASS: Component 2 — Found {procedure_step_tables} procedure step tables (0.30 pts)")
            total_score += 0.30
        elif procedure_step_tables > 0:
            print(f"PARTIAL: Component 2 — Found {procedure_step_tables}/10 procedure step tables ({step_score:.2f} pts)")
            total_score += step_score
        else:
            print(f"FAIL: Component 2 — No procedure step tables found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Spanish text is italic, English is not (0.20 pts)
    # Check across multiple bilingual 2-col tables:
    # - Right column (Spanish) runs should have italic=True
    # - Left column (English) runs should NOT have italic=True
    try:
        spanish_italic_count = 0
        spanish_total_count = 0
        english_not_italic_count = 0
        english_total_count = 0

        for t in tables:
            try:
                num_cols = len(t.columns)
                if num_cols != 2:
                    continue
                # Check header to confirm it's a bilingual English/Espanol table
                header_left = t.cell(0, 0).text.strip().lower()
                header_right = t.cell(0, 1).text.strip().lower()
                if not ('english' in header_left or 'step' in header_left.lower()):
                    # Also check if left col 0 has "English" header
                    if 'english' not in header_left:
                        continue

                for ri, row in enumerate(t.rows):
                    if ri == 0:
                        continue  # skip header
                    # Left column (English)
                    for para in row.cells[0].paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                english_total_count += 1
                                if not run.font.italic:
                                    english_not_italic_count += 1
                    # Right column (Spanish)
                    for para in row.cells[1].paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                spanish_total_count += 1
                                if run.font.italic:
                                    spanish_italic_count += 1
            except Exception:
                continue

        if spanish_total_count > 0 and english_total_count > 0:
            spanish_italic_ratio = spanish_italic_count / spanish_total_count
            english_regular_ratio = english_not_italic_count / english_total_count

            # Both ratios should be >= 0.8 for full credit
            if spanish_italic_ratio >= 0.8 and english_regular_ratio >= 0.8:
                print(f"PASS: Component 3 — Spanish italic: {spanish_italic_count}/{spanish_total_count} "
                      f"({spanish_italic_ratio:.0%}), English regular: {english_not_italic_count}/{english_total_count} "
                      f"({english_regular_ratio:.0%}) (0.20 pts)")
                total_score += 0.20
            elif spanish_italic_ratio >= 0.5 or english_regular_ratio >= 0.5:
                partial = 0.10
                print(f"PARTIAL: Component 3 — Spanish italic: {spanish_italic_ratio:.0%}, "
                      f"English regular: {english_regular_ratio:.0%} ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 — Spanish italic: {spanish_italic_ratio:.0%}, "
                      f"English regular: {english_regular_ratio:.0%}")
        else:
            print(f"FAIL: Component 3 — Insufficient bilingual table runs "
                  f"(Spanish: {spanish_total_count}, English: {english_total_count})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Emergency contacts table (0.10 pts)
    # A table with 3+ columns, 8+ rows, containing contact info
    try:
        contacts_found = False
        for t in tables:
            try:
                num_cols = len(t.columns)
                num_rows = len(t.rows)
                if num_cols >= 3 and num_rows >= 8:
                    # Check for contact-related content
                    header_text = ' '.join(c.text.lower() for c in t.rows[0].cells)
                    if 'contact' in header_text or 'phone' in header_text or 'telefono' in header_text:
                        # Verify some known contacts exist
                        all_text = ' '.join(c.text for row in t.rows for c in row.cells)
                        if '911' in all_text and 'Alvarez' in all_text:
                            contacts_found = True
                            break
            except Exception:
                continue

        if contacts_found:
            print(f"PASS: Component 4 — Emergency contacts table found with bilingual headers (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 4 — No valid emergency contacts table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Safety equipment checklist table (0.10 pts)
    # A 2-column table with 10+ rows listing equipment and inspection schedules
    try:
        checklist_found = False
        for t in tables:
            try:
                num_cols = len(t.columns)
                num_rows = len(t.rows)
                if num_cols == 2 and num_rows >= 10:
                    # Check for equipment-related content
                    header_text = ' '.join(c.text.lower() for c in t.rows[0].cells)
                    if 'equip' in header_text or 'inspection' in header_text or 'inspeccion' in header_text:
                        # Verify some known equipment
                        all_text = ' '.join(c.text for row in t.rows for c in row.cells)
                        if 'Hard Hat' in all_text or 'Casco' in all_text:
                            checklist_found = True
                            break
            except Exception:
                continue

        if checklist_found:
            print(f"PASS: Component 5 — Safety equipment checklist table found (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 — No valid safety equipment checklist table found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # Component 6: Emergency evacuation map placeholder in a table (0.05 pts)
    # Golden has a single-cell table (table 27) containing bilingual placeholder text
    # Initial has this as plain paragraph text, so we require it to be INSIDE a table
    try:
        map_placeholder_found = False
        for t in tables:
            try:
                all_text = ' '.join(c.text for row in t.rows for c in row.cells).upper()
                if 'EVACUATION MAP' in all_text and 'MAPA DE EVACUACION' in all_text:
                    map_placeholder_found = True
                    break
            except Exception:
                continue

        if map_placeholder_found:
            print(f"PASS: Component 6 — Emergency evacuation map placeholder found (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 6 — No evacuation map placeholder found")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
