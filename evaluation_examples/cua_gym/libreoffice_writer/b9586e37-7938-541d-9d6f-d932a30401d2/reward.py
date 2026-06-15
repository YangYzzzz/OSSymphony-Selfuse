"""
Reward Script: Remove all italic formatting while preserving other formatting
Task ID: writer_txtfmt_057
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All italic formatting removed — no run has italic=True
  Component 2 (0.3 pts): Italic removed AND bold/underline/color formatting preserved
                          on the specific runs that had bold+italic / underline+italic / color+italic
  Component 3 (0.2 pts): Italic removed AND highlight formatting preserved on
                          the specific para-6 runs that had italic+highlight
"""

import os
from docx import Document
from docx.shared import RGBColor

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_txtfmt_057'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Task: Remove all italic formatting while preserving bold, underline, colors, highlighting.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Sanity check: document must have expected structure
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Document has too few paragraphs ({len(doc.paragraphs)}), likely corrupted.")
        print("REWARD: 0.0")
        return 0.0

    paras = doc.paragraphs

    # Component 1: All italic formatting removed (0.5 points)
    # In the initial state, 13 runs had italic=True spread across paragraphs 1-6.
    # The task requires ALL italic to be removed, so no run may have italic=True.
    # This FAILS on initial (13 italic runs present) and PASSES on golden (0 italic runs).
    try:
        italic_runs_found = []
        for i, para in enumerate(paras):
            for j, run in enumerate(para.runs):
                if run.font.italic is True:
                    italic_runs_found.append((i, j, run.text[:40]))

        if not italic_runs_found:
            print(f"PASS: Component 1 — No italic runs found in document (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Found {len(italic_runs_found)} run(s) still with italic=True:")
            for para_idx, run_idx, text in italic_runs_found[:5]:
                print(f"  Para {para_idx}, Run {run_idx}: '{text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Italic removed AND bold/underline/color preserved on affected runs (0.3 points)
    # These runs had italic=True in initial AND also had other formatting (bold, underline, or color).
    # We verify that:
    #   (a) italic is now False/None (not True) — the change
    #   (b) the OTHER formatting attribute is still present — preservation
    # This compound check FAILS on initial (italic=True there) and PASSES on golden.
    #
    # Affected runs (initial state had italic=True + other formatting):
    #   Para 1, Run 1 ('remarkable transformations'): bold=True + italic=True → need italic gone, bold kept
    #   Para 1, Run 3 ('indispensable tools'):         bold=True + italic=True → need italic gone, bold kept
    #   Para 2, Run 1 ('scale their operations'):      underline=True + italic=True → need italic gone, underline kept
    #   Para 2, Run 3 ('hybrid and multi-cloud'):      underline=True + italic=True → need italic gone, underline kept
    #   Para 3, Run 1 ('Ransomware attacks...'):       color=FF0000 + italic=True → need italic gone, color kept
    #   Para 3, Run 3 ('zero-trust security...'):      color=FF0000 + italic=True → need italic gone, color kept
    #   Para 4, Run 1 ('Early quantum advantage...'):  bold=True + color=0000FF + italic=True → need italic gone, bold+color kept
    #   Para 4, Run 3 ('practical quantum supremacy'): bold=True + color=0000FF + italic=True → need italic gone, bold+color kept
    try:
        comp2_checks = []
        red = RGBColor(0xFF, 0x00, 0x00)
        blue = RGBColor(0x00, 0x00, 0xFF)

        # Para 1, Run 1: bold=True AND italic NOT True
        if len(paras) > 1 and len(paras[1].runs) > 1:
            run = paras[1].runs[1]
            italic_removed = (run.font.italic is not True)
            bold_kept = (run.font.bold is True)
            comp2_checks.append(("para1_run1_bold+no-italic", italic_removed and bold_kept,
                                  f"italic={run.font.italic}, bold={run.font.bold}"))

        # Para 1, Run 3: bold=True AND italic NOT True
        if len(paras) > 1 and len(paras[1].runs) > 3:
            run = paras[1].runs[3]
            italic_removed = (run.font.italic is not True)
            bold_kept = (run.font.bold is True)
            comp2_checks.append(("para1_run3_bold+no-italic", italic_removed and bold_kept,
                                  f"italic={run.font.italic}, bold={run.font.bold}"))

        # Para 2, Run 1: underline=True AND italic NOT True
        if len(paras) > 2 and len(paras[2].runs) > 1:
            run = paras[2].runs[1]
            italic_removed = (run.font.italic is not True)
            underline_kept = (run.font.underline is True)
            comp2_checks.append(("para2_run1_underline+no-italic", italic_removed and underline_kept,
                                  f"italic={run.font.italic}, underline={run.font.underline}"))

        # Para 2, Run 3: underline=True AND italic NOT True
        if len(paras) > 2 and len(paras[2].runs) > 3:
            run = paras[2].runs[3]
            italic_removed = (run.font.italic is not True)
            underline_kept = (run.font.underline is True)
            comp2_checks.append(("para2_run3_underline+no-italic", italic_removed and underline_kept,
                                  f"italic={run.font.italic}, underline={run.font.underline}"))

        # Para 3, Run 1: color=FF0000 AND italic NOT True
        if len(paras) > 3 and len(paras[3].runs) > 1:
            run = paras[3].runs[1]
            italic_removed = (run.font.italic is not True)
            rgb = run.font.color.rgb if run.font.color and run.font.color.type else None
            color_kept = (rgb == red)
            comp2_checks.append(("para3_run1_red+no-italic", italic_removed and color_kept,
                                  f"italic={run.font.italic}, color={rgb}"))

        # Para 3, Run 3: color=FF0000 AND italic NOT True
        if len(paras) > 3 and len(paras[3].runs) > 3:
            run = paras[3].runs[3]
            italic_removed = (run.font.italic is not True)
            rgb = run.font.color.rgb if run.font.color and run.font.color.type else None
            color_kept = (rgb == red)
            comp2_checks.append(("para3_run3_red+no-italic", italic_removed and color_kept,
                                  f"italic={run.font.italic}, color={rgb}"))

        # Para 4, Run 1: bold=True AND color=0000FF AND italic NOT True
        if len(paras) > 4 and len(paras[4].runs) > 1:
            run = paras[4].runs[1]
            italic_removed = (run.font.italic is not True)
            bold_kept = (run.font.bold is True)
            rgb = run.font.color.rgb if run.font.color and run.font.color.type else None
            color_kept = (rgb == blue)
            comp2_checks.append(("para4_run1_bold+blue+no-italic", italic_removed and bold_kept and color_kept,
                                  f"italic={run.font.italic}, bold={run.font.bold}, color={rgb}"))

        # Para 4, Run 3: bold=True AND color=0000FF AND italic NOT True
        if len(paras) > 4 and len(paras[4].runs) > 3:
            run = paras[4].runs[3]
            italic_removed = (run.font.italic is not True)
            bold_kept = (run.font.bold is True)
            rgb = run.font.color.rgb if run.font.color and run.font.color.type else None
            color_kept = (rgb == blue)
            comp2_checks.append(("para4_run3_bold+blue+no-italic", italic_removed and bold_kept and color_kept,
                                  f"italic={run.font.italic}, bold={run.font.bold}, color={rgb}"))

        passed = sum(1 for _, v, _ in comp2_checks if v)
        total_checks = len(comp2_checks)

        if total_checks > 0 and passed == total_checks:
            print(f"PASS: Component 2 — All {total_checks} bold/underline/color-with-italic-removed checks passed (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 — Only {passed}/{total_checks} checks passed:")
            for name, result, details in comp2_checks:
                status = "PASS" if result else "FAIL"
                print(f"  {status}: {name}: {details}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Italic removed AND highlight preserved on para 6 runs (0.2 points)
    # Para 6, Runs 1 and 3 had italic=True AND highlight=YELLOW in initial state.
    # The task requires italic removed while preserving highlight.
    # We verify: italic is NOT True AND highlight is still YELLOW.
    # This compound check FAILS on initial (italic=True) and PASSES on golden.
    try:
        from docx.enum.text import WD_COLOR_INDEX
        YELLOW = WD_COLOR_INDEX.YELLOW  # value 7

        comp3_checks = []
        if len(paras) > 6:
            p6_runs = paras[6].runs

            # Run 1: italic NOT True AND highlight=YELLOW
            if len(p6_runs) > 1:
                run = p6_runs[1]
                italic_removed = (run.font.italic is not True)
                highlight_kept = (run.font.highlight_color == YELLOW)
                comp3_checks.append(("para6_run1_yellow+no-italic", italic_removed and highlight_kept,
                                      f"italic={run.font.italic}, highlight={run.font.highlight_color}"))

            # Run 3: italic NOT True AND highlight=YELLOW
            if len(p6_runs) > 3:
                run = p6_runs[3]
                italic_removed = (run.font.italic is not True)
                highlight_kept = (run.font.highlight_color == YELLOW)
                comp3_checks.append(("para6_run3_yellow+no-italic", italic_removed and highlight_kept,
                                      f"italic={run.font.italic}, highlight={run.font.highlight_color}"))

        passed_c3 = sum(1 for _, v, _ in comp3_checks if v)
        total_c3 = len(comp3_checks)

        if total_c3 > 0 and passed_c3 == total_c3:
            print(f"PASS: Component 3 — {passed_c3}/{total_c3} highlight+no-italic checks passed (0.2 pts)")
            total_score += 0.2
        elif total_c3 > 0 and passed_c3 == 1:
            print(f"PARTIAL: Component 3 — {passed_c3}/{total_c3} highlight+no-italic checks passed (0.1 pts)")
            total_score += 0.1
        else:
            print(f"FAIL: Component 3 — {passed_c3}/{total_c3} checks passed:")
            for name, result, details in comp3_checks:
                status = "PASS" if result else "FAIL"
                print(f"  {status}: {name}: {details}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/review_article.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
