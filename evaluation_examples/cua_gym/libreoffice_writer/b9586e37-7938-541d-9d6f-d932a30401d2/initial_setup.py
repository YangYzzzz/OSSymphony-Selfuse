"""
Initial Setup: Remove all italic formatting while preserving other formatting
Task ID: writer_txtfmt_057
Domain: libreoffice_writer

Creates review_article.docx on ~/Desktop/ with a title in bold 16pt Arial
and six paragraphs containing complex mixed formatting including italics.
The agent's task is to remove all italic formatting while preserving everything else.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_txtfmt_057'
FILENAME = 'review_article.docx'
OUTPUT = f'{WORKDIR}/{FILENAME}'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    # Bold, 16pt Arial title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Annual Technology Trends: A Comprehensive Review")
    title_run.bold = True
    title_run.italic = False
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)

    # ── Paragraph 1: bold + italic words ──────────────────────────────────────
    p1 = doc.add_paragraph()
    r1a = p1.add_run("The technology landscape has undergone ")
    r1a.bold = False
    r1a.italic = False

    r1b = p1.add_run("remarkable transformations")
    r1b.bold = True
    r1b.italic = True   # bold+italic

    r1c = p1.add_run(" over the past decade. Artificial intelligence and machine learning have become ")
    r1c.bold = False
    r1c.italic = False

    r1d = p1.add_run("indispensable tools")
    r1d.bold = True
    r1d.italic = True   # bold+italic

    r1e = p1.add_run(" in virtually every industry, driving innovation at an unprecedented pace.")
    r1e.bold = False
    r1e.italic = False

    # ── Paragraph 2: italic + underline ────────────────────────────────────────
    p2 = doc.add_paragraph()
    r2a = p2.add_run("Cloud computing infrastructure has enabled organizations to ")
    r2a.italic = False
    r2a.underline = False

    r2b = p2.add_run("scale their operations globally")
    r2b.italic = True
    r2b.underline = True   # italic+underline

    r2c = p2.add_run(" without significant upfront capital investment. The shift from on-premises to ")

    r2d = p2.add_run("hybrid and multi-cloud deployments")
    r2d.italic = True
    r2d.underline = True   # italic+underline

    r2e = p2.add_run(" represents one of the most significant architectural changes in recent memory.")

    # ── Paragraph 3: italic + red color ────────────────────────────────────────
    p3 = doc.add_paragraph()
    r3a = p3.add_run("Cybersecurity threats continue to evolve in sophistication and frequency. ")
    r3a.italic = False

    r3b = p3.add_run("Ransomware attacks and data breaches")
    r3b.italic = True
    r3b.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)   # italic+red

    r3c = p3.add_run(" have cost organizations billions of dollars annually. Experts warn that ")
    r3c.italic = False

    r3d = p3.add_run("zero-trust security frameworks")
    r3d.italic = True
    r3d.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)   # italic+red

    r3e = p3.add_run(" are no longer optional but essential for modern enterprises.")
    r3e.italic = False

    # ── Paragraph 4: bold + italic + blue color ────────────────────────────────
    p4 = doc.add_paragraph()
    r4a = p4.add_run("Quantum computing is poised to revolutionize cryptography and complex problem-solving. ")
    r4a.italic = False

    r4b = p4.add_run("Early quantum advantage demonstrations")
    r4b.bold = True
    r4b.italic = True
    r4b.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)   # bold+italic+blue

    r4c = p4.add_run(" have sparked debate about the timeline for ")

    r4d = p4.add_run("practical quantum supremacy")
    r4d.bold = True
    r4d.italic = True
    r4d.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)   # bold+italic+blue

    r4e = p4.add_run(" in commercially relevant scenarios. Investment in this field has grown exponentially.")
    r4e.italic = False

    # ── Paragraph 5: italic-only words ─────────────────────────────────────────
    p5 = doc.add_paragraph()
    r5a = p5.add_run("Sustainable technology practices are gaining prominence as organizations recognize their ")
    r5a.italic = False

    r5b = p5.add_run("environmental and social responsibilities")
    r5b.italic = True   # italic-only

    r5c = p5.add_run(". From ")

    r5d = p5.add_run("green data centers")
    r5d.italic = True   # italic-only

    r5e = p5.add_run(" powered by renewable energy to ")

    r5f = p5.add_run("circular economy initiatives")
    r5f.italic = True   # italic-only

    r5g = p5.add_run(", the industry is beginning to align profit motives with planetary health.")
    r5g.italic = False

    # ── Paragraph 6: italic + highlight ────────────────────────────────────────
    p6 = doc.add_paragraph()
    r6a = p6.add_run("The convergence of augmented reality, virtual reality, and the metaverse presents ")
    r6a.italic = False

    r6b = p6.add_run("unprecedented opportunities for immersive experiences")
    r6b.italic = True
    r6b.font.highlight_color = WD_COLOR_INDEX.YELLOW   # italic+highlight

    r6c = p6.add_run(". While mainstream adoption remains ")

    r6d = p6.add_run("several years away")
    r6d.italic = True
    r6d.font.highlight_color = WD_COLOR_INDEX.YELLOW   # italic+highlight

    r6e = p6.add_run(", the foundations being laid today will define the next era of human-computer interaction.")
    r6e.italic = False

    # Save document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
