"""
Initial Setup: Create policy_brief.docx for writer_txtfmt_018
Task ID: writer_txtfmt_018
Domain: libreoffice_writer

Creates a policy brief document with four section headings (Introduction,
Background, Analysis, Conclusion) in 14pt Calibri Bold black, each followed
by a body text paragraph. Headings must NOT have color #003366 applied yet.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_018'
OUTPUT = f'{WORKDIR}/policy_brief.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Define section content: (heading_text, body_text)
    sections = [
        (
            "Introduction",
            "This policy brief examines the current regulatory landscape affecting "
            "small and medium enterprises (SMEs) in the manufacturing sector. Over the "
            "past decade, compliance costs have risen by an estimated 34%, placing "
            "disproportionate burdens on businesses with fewer than 250 employees. "
            "The analysis draws on survey data from 1,200 firms across 18 jurisdictions "
            "collected between January and September 2024."
        ),
        (
            "Background",
            "The regulatory environment for SMEs has evolved substantially since the "
            "passage of the Enterprise Reform Act of 2018. Prior to this legislation, "
            "manufacturing firms faced a fragmented patchwork of local, regional, and "
            "federal requirements that often conflicted with one another. Industry groups "
            "such as the National Federation of Independent Businesses (NFIB) reported "
            "that member firms spent an average of 2,340 hours annually on compliance "
            "activities — time that could otherwise be dedicated to innovation and growth."
        ),
        (
            "Analysis",
            "Our analysis reveals three key findings. First, compliance expenditure as "
            "a share of revenue is inversely correlated with firm size (r = -0.71, "
            "p < 0.001), confirming that smaller firms bear a heavier relative burden. "
            "Second, the introduction of unified digital reporting portals reduced "
            "average filing time by 28% in pilot jurisdictions. Third, sector-specific "
            "exemptions introduced in 2021 have yielded estimated savings of $4.2 billion "
            "for qualifying firms, though uptake remains uneven across regions."
        ),
        (
            "Conclusion",
            "The evidence supports a targeted reform agenda focused on three pillars: "
            "harmonisation of reporting standards across jurisdictions, expansion of the "
            "digital portal program to all eligible firms, and revision of the exemption "
            "thresholds to better reflect current economic conditions. Policymakers should "
            "prioritise consultation with SME representatives during the upcoming review "
            "cycle scheduled for Q2 2025. Failure to act risks further erosion of "
            "manufacturing capacity and continued displacement of domestic production."
        ),
    ]

    for heading_text, body_text in sections:
        # Add heading paragraph manually to control font precisely
        heading_para = doc.add_paragraph()
        run = heading_para.add_run(heading_text)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        # Explicitly set color to black (no dark blue — that is the task to do)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Add body paragraph
        body_para = doc.add_paragraph()
        body_run = body_para.add_run(body_text)
        body_run.font.name = "Calibri"
        body_run.font.size = Pt(11)
        body_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
