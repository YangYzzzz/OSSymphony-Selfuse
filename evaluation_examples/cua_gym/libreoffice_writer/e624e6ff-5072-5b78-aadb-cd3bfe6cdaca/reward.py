"""
Reward Script: Configure different headers on odd and even pages
Task ID: writer_biz_075
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): evenAndOddHeaders enabled in settings.xml
  Component 2 (0.35): Odd (default) header = 'Meridian Solutions Inc.' right-aligned
  Component 3 (0.35): Even header = 'Annual Report 2025' left-aligned
  Component 4 (0.05): Odd header does NOT contain 'Annual Report 2025' (exclusivity)
"""

import os
import zipfile
import io

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_075'


def persist_app_state(domain: str):
    """Attempt to save any unsaved changes in LibreOffice."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load the docx as a zip to inspect settings.xml
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        zf = zipfile.ZipFile(io.BytesIO(data))
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Load with python-docx for header inspection
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot parse docx {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: evenAndOddHeaders enabled in settings.xml (0.25 points)
    # This is the key toggle that enables different headers on odd/even pages.
    # Initial state: NOT present. Golden state: present.
    try:
        if 'word/settings.xml' in zf.namelist():
            from lxml import etree
            settings_xml = zf.read('word/settings.xml')
            root = etree.fromstring(settings_xml)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            eaoh = root.findall('.//w:evenAndOddHeaders', ns)
            if len(eaoh) > 0:
                print(f"PASS: Component 1 — evenAndOddHeaders is enabled in settings.xml (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 — evenAndOddHeaders NOT found in settings.xml")
        else:
            print(f"FAIL: Component 1 — word/settings.xml not found in docx")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Odd (default) header contains 'Meridian Solutions Inc.' right-aligned (0.35 points)
    # Initial state: header text is 'Meridian Solutions Inc. — Annual Report 2025' CENTER aligned.
    # Golden state: default header text is 'Meridian Solutions Inc.' RIGHT aligned.
    try:
        section = doc.sections[0]
        header = section.header
        header_text = ''
        header_alignment = None
        for p in header.paragraphs:
            if p.text.strip():
                header_text = p.text.strip()
                header_alignment = p.paragraph_format.alignment
                break

        text_match = 'Meridian Solutions Inc.' in header_text
        # Must be RIGHT aligned (enum value 2)
        right_aligned = (header_alignment is not None and
                         header_alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if text_match and right_aligned:
            print(f"PASS: Component 2 — Odd header = '{header_text}', alignment=RIGHT (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 2 — Odd header text='{header_text}' (match={text_match}), "
                  f"alignment={header_alignment} (right={right_aligned})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Even header contains 'Annual Report 2025' left-aligned (0.35 points)
    # Initial state: even header is empty. Golden state: 'Annual Report 2025' LEFT aligned.
    try:
        section = doc.sections[0]
        even_header = section.even_page_header
        even_text = ''
        even_alignment = None
        for p in even_header.paragraphs:
            if p.text.strip():
                even_text = p.text.strip()
                even_alignment = p.paragraph_format.alignment
                break

        text_match = 'Annual Report 2025' in even_text
        # Must be LEFT aligned (enum value 0) or None (default=LEFT)
        left_aligned = (even_alignment is None or
                        even_alignment == WD_PARAGRAPH_ALIGNMENT.LEFT)

        if text_match and left_aligned:
            print(f"PASS: Component 3 — Even header = '{even_text}', alignment=LEFT (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 3 — Even header text='{even_text}' (match={text_match}), "
                  f"alignment={even_alignment} (left={left_aligned})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Odd header does NOT contain 'Annual Report 2025' (0.05 points)
    # This ensures the headers were properly split, not just duplicated.
    # Initial state: odd header contains BOTH texts. Golden state: only company name.
    try:
        section = doc.sections[0]
        header = section.header
        odd_full_text = ' '.join(p.text for p in header.paragraphs).strip()

        if 'Annual Report 2025' not in odd_full_text:
            print(f"PASS: Component 4 — Odd header correctly excludes 'Annual Report 2025' (0.05 pts)")
            total_score += 0.05
        else:
            print(f"FAIL: Component 4 — Odd header still contains 'Annual Report 2025': '{odd_full_text}'")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
