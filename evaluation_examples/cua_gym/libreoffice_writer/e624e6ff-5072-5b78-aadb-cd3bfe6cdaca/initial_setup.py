"""
Initial Setup: Configure different headers on odd and even pages
Task ID: writer_biz_075
Domain: libreoffice_writer

Creates a 12-page Writer document with realistic Annual Report 2025 content
for Meridian Solutions Inc. All pages have the SAME header (default behavior).
The task is for the agent to configure different odd/even page headers.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_075'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Same header on all pages (default behavior) ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Meridian Solutions Inc. — Annual Report 2025"
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in hp.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Footer with page numbers ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r0 = fp.add_run("Page ")
    r0.font.size = Pt(9)
    r1 = fp.add_run()
    r1._element.append(r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    r2._element.append(instr)
    r3 = fp.add_run()
    r3._element.append(r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    # ===== Page 1: Title Page =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Meridian Solutions Inc.")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run("Annual Report 2025")
    run2.font.size = Pt(22)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p3.paragraph_format.space_before = Pt(40)
    run3 = p3.add_run("Presented to Shareholders and Stakeholders\nFiscal Year Ending December 31, 2025")
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ===== Page 2: Table of Contents =====
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("1. Executive Summary", "3"),
        ("2. Financial Highlights", "4"),
        ("3. Revenue Breakdown by Segment", "5"),
        ("4. Operational Performance", "6"),
        ("5. Strategic Initiatives", "7"),
        ("6. Human Capital Report", "8"),
        ("7. Sustainability & ESG", "9"),
        ("8. Risk Management", "10"),
        ("9. Market Outlook", "11"),
        ("10. Board of Directors", "12"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run_item = p.add_run(f"{item}")
        run_item.font.size = Pt(11)
        run_dots = p.add_run("." * (60 - len(item)))
        run_dots.font.size = Pt(11)
        run_dots.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run_pg = p.add_run(f" {page}")
        run_pg.font.size = Pt(11)

    # ===== Page 3: Executive Summary =====
    doc.add_page_break()
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Dear Shareholders and Stakeholders,\n\n"
        "Fiscal year 2025 has been a transformative period for Meridian Solutions Inc. "
        "Our consolidated revenue reached $487.3 million, reflecting a 14.2% year-over-year increase "
        "driven by strong performance across our cloud services and enterprise consulting divisions. "
        "Operating income improved to $68.9 million, representing a 14.1% operating margin."
    )
    doc.add_paragraph(
        "Our strategic investments in artificial intelligence and machine learning capabilities "
        "have begun to yield measurable returns. The launch of our MeridianAI platform in Q2 "
        "attracted 340 enterprise clients within the first six months, generating $23.7 million "
        "in incremental annual recurring revenue."
    )
    doc.add_paragraph(
        "We successfully completed the acquisition of DataBridge Analytics in September 2025, "
        "expanding our data engineering capabilities and adding 180 technical professionals "
        "to our workforce. This acquisition strengthens our competitive position in the "
        "rapidly growing data infrastructure market."
    )

    # ===== Page 4: Financial Highlights =====
    doc.add_page_break()
    doc.add_heading("2. Financial Highlights", level=1)
    doc.add_paragraph(
        "The following table summarizes our key financial metrics for fiscal year 2025 "
        "compared to the prior year:"
    )
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    headers = ["Metric", "FY 2025", "FY 2024"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data = [
        ["Total Revenue", "$487.3M", "$426.8M"],
        ["Operating Income", "$68.9M", "$55.2M"],
        ["Net Income", "$52.1M", "$41.6M"],
        ["Earnings Per Share", "$3.47", "$2.78"],
        ["Free Cash Flow", "$74.5M", "$62.3M"],
        ["Total Assets", "$892.4M", "$745.1M"],
        ["Employees", "4,280", "3,760"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # ===== Page 5: Revenue Breakdown =====
    doc.add_page_break()
    doc.add_heading("3. Revenue Breakdown by Segment", level=1)
    doc.add_paragraph(
        "Our diversified revenue streams provide resilience and growth opportunities "
        "across multiple market segments:"
    )
    segments = [
        ("Cloud Infrastructure Services", "$178.4M", "36.6%", "+18.3%"),
        ("Enterprise Consulting", "$132.6M", "27.2%", "+11.7%"),
        ("Data & Analytics Platform", "$89.2M", "18.3%", "+22.4%"),
        ("Managed IT Services", "$54.8M", "11.2%", "+5.1%"),
        ("Cybersecurity Solutions", "$32.3M", "6.6%", "+28.9%"),
    ]
    table2 = doc.add_table(rows=len(segments) + 1, cols=4)
    table2.style = "Table Grid"
    for i, h in enumerate(["Segment", "Revenue", "% of Total", "YoY Growth"]):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r, seg in enumerate(segments, 1):
        for c, val in enumerate(seg):
            table2.cell(r, c).text = val

    # ===== Page 6: Operational Performance =====
    doc.add_page_break()
    doc.add_heading("4. Operational Performance", level=1)
    doc.add_paragraph(
        "Operational efficiency remained a top priority in 2025. Our global delivery centers "
        "in Austin, Toronto, Bangalore, and Dublin maintained a combined utilization rate of 82.3%, "
        "up from 78.9% in the prior year."
    )
    doc.add_paragraph(
        "Key operational achievements include:\n"
    )
    bullets = [
        "Reduced average project delivery time by 17% through agile methodology adoption",
        "Achieved 99.97% uptime across managed cloud environments",
        "Implemented automated testing pipelines reducing QA cycle time by 34%",
        "Expanded Austin data center capacity by 40% to meet growing demand",
        "Achieved SOC 2 Type II and ISO 27001 recertification",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ===== Page 7: Strategic Initiatives =====
    doc.add_page_break()
    doc.add_heading("5. Strategic Initiatives", level=1)
    doc.add_paragraph(
        "Our three-year strategic roadmap, launched in early 2024, centers on three pillars: "
        "innovation acceleration, geographic expansion, and talent development."
    )
    doc.add_heading("5.1 MeridianAI Platform", level=2)
    doc.add_paragraph(
        "The MeridianAI platform leverages proprietary machine learning models to help enterprise "
        "clients automate document processing, customer service workflows, and predictive maintenance "
        "operations. In its first year, the platform processed over 12 million documents and reduced "
        "client processing costs by an average of 43%."
    )
    doc.add_heading("5.2 Asia-Pacific Expansion", level=2)
    doc.add_paragraph(
        "We opened new offices in Singapore and Tokyo in 2025, establishing a direct presence "
        "in the Asia-Pacific market. These offices already serve 28 enterprise clients and "
        "contributed $14.2 million in revenue during their partial year of operations."
    )

    # ===== Page 8: Human Capital =====
    doc.add_page_break()
    doc.add_heading("6. Human Capital Report", level=1)
    doc.add_paragraph(
        "Our people are the foundation of our success. As of December 31, 2025, Meridian Solutions "
        "employs 4,280 professionals across 12 offices worldwide."
    )
    doc.add_paragraph(
        "Employee satisfaction score: 4.3 out of 5.0 (annual engagement survey)\n"
        "Voluntary turnover rate: 11.2% (industry average: 15.8%)\n"
        "Average training hours per employee: 48 hours\n"
        "Women in leadership positions: 38% (up from 33% in 2024)\n"
        "Internal promotion rate: 24%"
    )
    doc.add_paragraph(
        "We invested $8.4 million in employee development programs, including our flagship "
        "Meridian Academy technical certification program, which graduated 620 professionals "
        "in cloud architecture, data engineering, and cybersecurity disciplines."
    )

    # ===== Page 9: Sustainability =====
    doc.add_page_break()
    doc.add_heading("7. Sustainability & ESG", level=1)
    doc.add_paragraph(
        "Meridian Solutions is committed to responsible business practices and environmental stewardship. "
        "Our 2025 ESG report highlights the following achievements:"
    )
    esg_items = [
        "Reduced Scope 1 and 2 carbon emissions by 22% through renewable energy procurement",
        "100% of data centers powered by renewable energy sources",
        "Launched the Meridian Green IT Advisory practice to help clients reduce their technology carbon footprint",
        "Donated $2.1 million to STEM education programs in underserved communities",
        "Achieved carbon neutral certification for all office operations",
        "Implemented paperless workflows across all internal business processes",
    ]
    for item in esg_items:
        doc.add_paragraph(item, style="List Bullet")

    # ===== Page 10: Risk Management =====
    doc.add_page_break()
    doc.add_heading("8. Risk Management", level=1)
    doc.add_paragraph(
        "Our enterprise risk management framework identifies, assesses, and mitigates risks "
        "across four categories: operational, financial, regulatory, and reputational."
    )
    doc.add_paragraph(
        "Key risk factors and mitigation strategies for 2026 include:"
    )
    risks = [
        ("Cybersecurity Threats", "Continuous investment in zero-trust architecture and 24/7 SOC operations"),
        ("Talent Acquisition", "Competitive compensation packages, remote work options, and career development pathways"),
        ("Regulatory Compliance", "Dedicated compliance team monitoring evolving data privacy regulations across 15 jurisdictions"),
        ("Economic Uncertainty", "Diversified client base across industries and geographies reducing concentration risk"),
    ]
    risk_table = doc.add_table(rows=len(risks) + 1, cols=2)
    risk_table.style = "Table Grid"
    risk_table.cell(0, 0).text = "Risk Category"
    risk_table.cell(0, 1).text = "Mitigation Strategy"
    for run in risk_table.cell(0, 0).paragraphs[0].runs:
        run.bold = True
    for run in risk_table.cell(0, 1).paragraphs[0].runs:
        run.bold = True
    for r, (risk, mitigation) in enumerate(risks, 1):
        risk_table.cell(r, 0).text = risk
        risk_table.cell(r, 1).text = mitigation

    # ===== Page 11: Market Outlook =====
    doc.add_page_break()
    doc.add_heading("9. Market Outlook", level=1)
    doc.add_paragraph(
        "The global IT services market is projected to reach $1.8 trillion by 2027, growing at a "
        "CAGR of 8.4%. Meridian Solutions is well-positioned to capture this growth through our "
        "differentiated capabilities in cloud, AI, and cybersecurity."
    )
    doc.add_paragraph(
        "For fiscal year 2026, we provide the following guidance:\n"
    )
    doc.add_paragraph("Revenue: $545 million to $565 million (12-16% growth)", style="List Bullet")
    doc.add_paragraph("Operating margin: 14.5% to 15.5%", style="List Bullet")
    doc.add_paragraph("Earnings per share: $3.85 to $4.10", style="List Bullet")
    doc.add_paragraph("Capital expenditure: $35 million to $40 million", style="List Bullet")
    doc.add_paragraph(
        "We remain confident in our ability to deliver sustainable long-term value for our "
        "shareholders while investing in the capabilities and talent needed to address our "
        "clients' most pressing technology challenges."
    )

    # ===== Page 12: Board of Directors =====
    doc.add_page_break()
    doc.add_heading("10. Board of Directors", level=1)
    board_members = [
        ("Patricia Hayward", "Chairperson", "Former CEO, Vertex Technologies; Board member since 2018"),
        ("Robert Chen", "CEO & Director", "Co-founder, Meridian Solutions; Leading since 2012"),
        ("Dr. Angela Martinez", "Independent Director", "Professor of Computer Science, MIT; AI ethics advisor"),
        ("James Okonkwo", "Independent Director", "Former CFO, GlobalTech Partners; Audit committee chair"),
        ("Sarah Lindberg", "Independent Director", "Managing Partner, Nordic Ventures; ESG committee chair"),
        ("David Nakamura", "Independent Director", "Former CTO, Pacific Digital; Technology committee chair"),
    ]
    board_table = doc.add_table(rows=len(board_members) + 1, cols=3)
    board_table.style = "Table Grid"
    for i, h in enumerate(["Name", "Role", "Background"]):
        cell = board_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r, (name, role, bg) in enumerate(board_members, 1):
        board_table.cell(r, 0).text = name
        board_table.cell(r, 1).text = role
        board_table.cell(r, 2).text = bg

    doc.add_paragraph()
    p_closing = doc.add_paragraph()
    p_closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_closing.paragraph_format.space_before = Pt(24)
    run_closing = p_closing.add_run(
        "This report was prepared by the Office of Investor Relations.\n"
        "Meridian Solutions Inc. | 2100 Innovation Blvd, Austin, TX 78701\n"
        "investor.relations@meridiansolutions.com | (512) 555-0147"
    )
    run_closing.font.size = Pt(9)
    run_closing.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
