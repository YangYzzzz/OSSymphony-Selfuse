"""
Initial Setup: Create a Writer document with encyclopedia content and marked index entries
Task ID: writer_af_038
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_af_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_xe_field(paragraph, entry_text):
    """Add an XE (index entry) field to a paragraph using Word XML field codes."""
    # Begin field char
    r1 = paragraph.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    # Instruction text
    r2 = paragraph.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = f' XE "{entry_text}" '
    r2._element.append(instr)

    # End field char
    r3 = paragraph.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading('The Comprehensive Encyclopedia of Natural Sciences', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Reference Guide to the Foundations of Scientific Knowledge')
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # 20 index entry terms spread across content
    index_terms = [
        "Photosynthesis", "Mitochondria", "Tectonic Plates", "Electromagnetic Spectrum",
        "Quantum Mechanics", "Natural Selection", "Thermodynamics", "Biosynthesis",
        "Cellular Respiration", "Gravitational Waves", "Genetic Drift", "Ecosystem Dynamics",
        "Nuclear Fusion", "Atmospheric Chemistry", "Plate Tectonics", "Stellar Evolution",
        "Molecular Biology", "Continental Drift", "Organic Chemistry", "Relativistic Physics"
    ]

    # Encyclopedia content organized into sections
    sections_content = [
        {
            "title": "Chapter 1: The Foundations of Biology",
            "paragraphs": [
                "Biology is the scientific study of life, encompassing a vast array of disciplines that explore the structure, function, growth, origin, evolution, and distribution of living organisms. From the smallest microorganism to the largest whale, biology seeks to understand the mechanisms that drive life on our planet. The field has expanded dramatically since the early observations of naturalists in the 17th and 18th centuries, and today it incorporates advanced technologies such as genomic sequencing, bioinformatics, and molecular imaging.",
                "At the heart of every living cell lies the process of {Photosynthesis}, through which plants and certain bacteria convert sunlight into chemical energy. This remarkable biochemical pathway involves the absorption of light by chlorophyll molecules in the thylakoid membranes, the splitting of water molecules, and the subsequent fixation of carbon dioxide into glucose. Without photosynthesis, the oxygen-rich atmosphere we depend on would not exist, and nearly all food chains on Earth would collapse.",
                "The {Mitochondria} are often described as the powerhouses of the cell, responsible for generating the majority of adenosine triphosphate (ATP) used as a source of chemical energy. These double-membraned organelles possess their own circular DNA, lending support to the endosymbiotic theory, which proposes that mitochondria were once free-living prokaryotes engulfed by ancestral eukaryotic cells. The inner membrane of the mitochondrion is folded into cristae, which increase the surface area available for oxidative phosphorylation.",
                "The phenomenon of {Natural Selection}, first articulated by Charles Darwin and Alfred Russel Wallace, remains one of the most powerful explanatory frameworks in biology. It posits that individuals within a population exhibit variation in heritable traits, and those traits that confer a survival or reproductive advantage tend to become more common over successive generations. This differential reproductive success drives adaptation and, over geological timescales, speciation.",
                "{Cellular Respiration} is the metabolic process by which cells break down glucose and other organic molecules to produce ATP. This complex sequence of reactions occurs in three main stages: glycolysis in the cytoplasm, the citric acid cycle in the mitochondrial matrix, and the electron transport chain along the inner mitochondrial membrane. Together, these stages can yield up to 36-38 molecules of ATP per molecule of glucose oxidized.",
                "{Genetic Drift} refers to random fluctuations in allele frequencies within a population, particularly significant in small populations where chance events can dramatically alter the genetic composition. Unlike natural selection, genetic drift is not driven by the adaptive value of traits but rather by stochastic sampling effects during reproduction. The founder effect and population bottlenecks are classic examples of scenarios where genetic drift plays a prominent role.",
            ]
        },
        {
            "title": "Chapter 2: Earth Sciences and Geology",
            "paragraphs": [
                "The Earth sciences encompass the study of our planet's physical structure, its atmosphere, its oceans, and the dynamic processes that shape its surface. Geologists, meteorologists, oceanographers, and environmental scientists work together to understand the complex systems that govern Earth's behavior over timescales ranging from seconds to billions of years.",
                "The theory of {Tectonic Plates} revolutionized our understanding of Earth's geology by explaining how the lithosphere is divided into a series of rigid plates that float atop the semi-fluid asthenosphere. These plates interact at their boundaries through convergent, divergent, and transform motions, giving rise to earthquakes, volcanic activity, and the formation of mountain ranges. The mid-Atlantic ridge, for instance, is a divergent boundary where new oceanic crust is continuously being formed.",
                "{Plate Tectonics} as a unifying theory was not fully accepted until the 1960s, when seafloor spreading was confirmed through magnetic striping patterns on the ocean floor. The theory explains the distribution of fossils across continents, the occurrence of similar rock formations on distant landmasses, and the ring of fire that encircles the Pacific Ocean. It remains the foundational framework for understanding large-scale geological phenomena.",
                "The concept of {Continental Drift}, first proposed by Alfred Wegener in 1912, suggested that the continents were once joined in a single supercontinent called Pangaea and have since drifted apart. Although Wegener's hypothesis was initially met with skepticism due to the lack of a plausible mechanism, the subsequent discovery of plate tectonics provided the driving force behind continental movement. Evidence from paleoclimate data, fossil distributions, and the geometric fit of continental shelves supports this theory.",
                "{Atmospheric Chemistry} examines the chemical composition of Earth's atmosphere and the reactions that occur within it. Key topics include the formation and destruction of ozone in the stratosphere, the role of greenhouse gases in regulating global temperature, and the chemistry of acid rain. Anthropogenic emissions of sulfur dioxide, nitrogen oxides, and volatile organic compounds have significantly altered atmospheric chemistry, contributing to climate change and air quality degradation.",
            ]
        },
        {
            "title": "Chapter 3: Physics and the Universe",
            "paragraphs": [
                "Physics is the fundamental science that seeks to understand the behavior of matter and energy at every scale, from subatomic particles to the vast expanse of the cosmos. Its laws govern the motion of celestial bodies, the propagation of light, the flow of electric current, and the interactions of elementary particles.",
                "The {Electromagnetic Spectrum} encompasses the full range of electromagnetic radiation, from radio waves with wavelengths measured in meters to gamma rays with wavelengths smaller than atomic nuclei. Visible light occupies only a narrow band within this spectrum, yet it is the primary means by which humans perceive the world. Each region of the spectrum has unique properties and applications: radio waves for communication, microwaves for heating, infrared for thermal imaging, ultraviolet for sterilization, X-rays for medical imaging, and gamma rays for cancer treatment.",
                "{Quantum Mechanics} emerged in the early 20th century as physicists grappled with phenomena that classical mechanics could not explain, such as blackbody radiation, the photoelectric effect, and atomic spectra. The theory introduces concepts like wave-particle duality, the uncertainty principle, and quantum entanglement, which challenge our intuitive understanding of reality. Schrödinger's wave equation and Heisenberg's matrix mechanics provide the mathematical framework for predicting the behavior of particles at the quantum scale.",
                "The laws of {Thermodynamics} govern the behavior of energy and its transformations. The first law states that energy cannot be created or destroyed, only converted from one form to another. The second law introduces the concept of entropy, asserting that natural processes tend to increase the disorder of a system. The third law establishes absolute zero as a theoretical limit that can never be reached. These principles underpin everything from engine design to the fate of the universe.",
                "{Gravitational Waves}, predicted by Einstein's general theory of relativity in 1915 and first directly detected by LIGO in 2015, are ripples in the fabric of spacetime caused by the acceleration of massive objects. The merger of two black holes or neutron stars produces gravitational waves that propagate outward at the speed of light. Their detection has opened a new window on the universe, enabling astronomers to observe cosmic events that are invisible to electromagnetic telescopes.",
                "{Nuclear Fusion} is the process by which lighter atomic nuclei combine to form heavier nuclei, releasing enormous amounts of energy. It is the power source of stars, including our Sun, where hydrogen nuclei fuse into helium under extreme temperatures and pressures. Scientists have long sought to harness fusion energy on Earth as a clean and virtually limitless power source, with projects like ITER aiming to demonstrate the feasibility of sustained fusion reactions.",
                "{Stellar Evolution} describes the life cycle of stars from their formation in molecular clouds to their eventual death as white dwarfs, neutron stars, or black holes. A star's mass determines its evolutionary path: low-mass stars like our Sun will eventually expand into red giants before shedding their outer layers to form planetary nebulae, while massive stars undergo dramatic supernova explosions. The elements forged in stellar cores and expelled in these events enrich the interstellar medium, providing the raw materials for new generations of stars and planets.",
                "{Relativistic Physics} extends classical mechanics to objects moving at speeds approaching the speed of light. Einstein's special theory of relativity, published in 1905, introduced the concepts of time dilation and length contraction, demonstrating that measurements of time and space are not absolute but depend on the observer's relative motion. The famous equation E=mc² revealed the equivalence of mass and energy, with profound implications for nuclear physics and cosmology.",
            ]
        },
        {
            "title": "Chapter 4: Chemistry and Molecular Science",
            "paragraphs": [
                "Chemistry is the science of matter, investigating the properties, composition, and transformations of substances at the molecular and atomic level. It bridges physics and biology, providing the molecular understanding necessary for advances in medicine, materials science, and environmental protection.",
                "{Biosynthesis} refers to the enzyme-catalyzed processes within living organisms that produce complex molecules from simpler precursors. Examples include the synthesis of amino acids, nucleotides, fatty acids, and secondary metabolites such as alkaloids and terpenes. These pathways are tightly regulated by feedback mechanisms and allosteric control, ensuring that cells produce the molecules they need in appropriate quantities. Understanding biosynthetic pathways has enabled the development of antibiotics, anticancer agents, and industrial bioproducts.",
                "{Ecosystem Dynamics} encompasses the study of energy flow and nutrient cycling within ecosystems. Producers capture solar energy through photosynthesis and convert it into biomass, which is then consumed by primary consumers and passed along trophic levels to secondary and tertiary consumers. Decomposers break down dead organic matter, returning nutrients to the soil. The balance of these processes determines ecosystem productivity, stability, and resilience to disturbance. Human activities such as deforestation, pollution, and climate change increasingly disrupt these natural cycles.",
                "{Molecular Biology} focuses on the structure and function of the macromolecules essential to life, particularly nucleic acids and proteins. The central dogma of molecular biology describes the flow of genetic information from DNA to RNA to protein, a process mediated by transcription and translation. Advances in techniques such as PCR, CRISPR-Cas9 gene editing, and next-generation sequencing have revolutionized our ability to study and manipulate genetic material, with applications ranging from gene therapy to forensic science.",
                "{Organic Chemistry} is the branch of chemistry concerned with the structure, properties, and reactions of carbon-containing compounds. Carbon's ability to form four covalent bonds and create stable chains, rings, and branched structures gives rise to an extraordinary diversity of organic molecules. Key reaction types include substitution, elimination, addition, and rearrangement reactions. Organic chemistry is foundational to the pharmaceutical, polymer, petrochemical, and food industries, and it provides the molecular basis for understanding biochemistry and metabolism.",
            ]
        },
        {
            "title": "Chapter 5: Interdisciplinary Perspectives",
            "paragraphs": [
                "Modern science increasingly recognizes that the most challenging questions lie at the intersections of traditional disciplines. Climate change, for instance, requires expertise in atmospheric chemistry, ocean physics, ecology, and economics. Similarly, the development of new materials draws on quantum mechanics, organic chemistry, and engineering principles.",
                "The convergence of biology and physics has given rise to biophysics, a field that applies physical methods to biological problems. From the folding of proteins to the mechanics of cell division, biophysics provides quantitative insights that complement traditional biological approaches. Techniques such as X-ray crystallography, cryo-electron microscopy, and single-molecule fluorescence have enabled researchers to visualize biological structures at atomic resolution.",
                "Environmental science integrates knowledge from geology, chemistry, biology, and social science to address pressing issues such as biodiversity loss, water pollution, and sustainable resource management. The concept of planetary boundaries, which defines safe operating limits for human activity, draws on ecosystem dynamics, atmospheric chemistry, and biogeochemical cycling to identify thresholds beyond which Earth's systems may undergo irreversible change.",
                "The search for extraterrestrial life combines astrobiology with planetary science, drawing on our understanding of extremophiles, organic chemistry, and stellar evolution. The discovery of exoplanets in habitable zones, the detection of organic molecules on Mars and in cometary material, and the exploration of subsurface oceans on moons like Europa and Enceladus have reinvigorated the quest to determine whether life exists beyond Earth.",
                "As our understanding of the natural world deepens, the connections between physics, chemistry, biology, and earth science become ever more apparent. The reductionist approach that drove much of 20th-century science is increasingly complemented by systems thinking, which seeks to understand how components interact to produce emergent properties. This holistic perspective is essential for tackling the complex, multifaceted challenges that define the 21st century.",
            ]
        },
    ]

    # Track which paragraph to insert XE fields into
    term_index = 0

    for sec_content in sections_content:
        doc.add_heading(sec_content["title"], level=1)

        for para_text in sec_content["paragraphs"]:
            # Check if this paragraph contains an index term marker
            import re
            term_match = re.search(r'\{(\w[\w\s]*)\}', para_text)

            if term_match and term_index < len(index_terms):
                term = term_match.group(1)
                # Replace {Term} with just Term in text
                clean_text = para_text.replace('{' + term + '}', term)
                para = doc.add_paragraph()
                run = para.add_run(clean_text)
                run.font.size = Pt(11)
                run.font.name = 'Liberation Serif'

                # Add XE field for this term
                add_xe_field(para, term)
                term_index += 1
            else:
                para = doc.add_paragraph()
                run = para.add_run(para_text)
                run.font.size = Pt(11)
                run.font.name = 'Liberation Serif'

            # Add spacing for readability
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = 1.15

        # Add page breaks between chapters to reach ~30 pages
        doc.add_page_break()

    # Add additional filler content to reach approximately 30 pages
    additional_sections = [
        ("Appendix A: Historical Timeline of Scientific Discovery", [
            "The history of science is a chronicle of human curiosity, persistence, and ingenuity. From the ancient Greek philosophers who first proposed atomic theories to the modern physicists who detected gravitational waves, each generation has built upon the discoveries of its predecessors.",
            "In the 3rd century BCE, Archimedes of Syracuse established principles of buoyancy and leverage that remain fundamental to engineering. Centuries later, the Islamic Golden Age produced scholars like Ibn al-Haytham, whose work on optics laid the groundwork for the modern scientific method. The European Renaissance brought Copernicus, Galileo, and Kepler, whose astronomical observations overturned the geocentric model of the universe.",
            "The 17th century saw Newton's formulation of the laws of motion and universal gravitation, while the 18th century brought Lavoisier's systematic approach to chemistry and Linnaeus's classification of living organisms. The 19th century witnessed Darwin's theory of evolution, Maxwell's unification of electricity and magnetism, and Mendeleev's periodic table.",
            "The 20th century was marked by revolutionary discoveries: Einstein's relativity, quantum mechanics, the structure of DNA, and plate tectonics. The century closed with the Human Genome Project and the detection of the cosmic microwave background radiation, confirming the Big Bang theory.",
            "As we progress through the 21st century, science continues to push the boundaries of knowledge. Gravitational wave astronomy, CRISPR gene editing, artificial intelligence, and the search for dark matter represent just a few of the frontiers that promise to reshape our understanding of the natural world.",
        ]),
        ("Appendix B: Glossary of Key Concepts", [
            "Adaptation: The process by which organisms become better suited to their environment through natural selection over generations. Adaptations can be structural, physiological, or behavioral, and they reflect the interplay between genetic variation and environmental pressures.",
            "Entropy: A measure of the disorder or randomness in a system. In thermodynamics, entropy always increases in isolated systems, driving the directionality of natural processes. In information theory, entropy quantifies the uncertainty associated with a random variable.",
            "Homeostasis: The tendency of biological systems to maintain stable internal conditions despite external fluctuations. Examples include temperature regulation in mammals, blood glucose control, and pH buffering in cellular environments.",
            "Isotope: Atoms of the same element that differ in the number of neutrons in their nuclei. Isotopes may be stable or radioactive, and they have applications in radiometric dating, medical imaging, and nuclear energy.",
            "Metabolism: The sum of all chemical reactions occurring within a living organism. Metabolism is divided into catabolism (breakdown of molecules for energy) and anabolism (synthesis of complex molecules from simpler precursors).",
            "Phenotype: The observable characteristics of an organism, resulting from the interaction of its genotype with the environment. Phenotypic traits include morphology, behavior, physiological properties, and biochemical composition.",
            "Wavelength: The distance between successive crests of a wave. In the context of electromagnetic radiation, wavelength determines the type of radiation: longer wavelengths correspond to radio waves, while shorter wavelengths correspond to gamma rays.",
        ]),
        ("Appendix C: Further Reading and Resources", [
            "For readers seeking to deepen their understanding of the topics covered in this encyclopedia, the following resources provide excellent starting points. Each has been selected for its accessibility, rigor, and comprehensiveness.",
            "In biology, 'Molecular Biology of the Cell' by Alberts et al. offers a thorough introduction to cellular and molecular biology, while 'The Selfish Gene' by Richard Dawkins provides an engaging exploration of evolutionary theory through the lens of gene-centered evolution.",
            "For physics enthusiasts, 'The Feynman Lectures on Physics' by Richard Feynman remains an unparalleled resource for understanding fundamental physical concepts. 'A Brief History of Time' by Stephen Hawking offers an accessible introduction to cosmology and the nature of spacetime.",
            "In chemistry, 'Organic Chemistry' by Clayden, Greeves, and Warren is widely regarded as one of the best textbooks for understanding organic reactions and mechanisms. 'The Disappearing Spoon' by Sam Kean provides an entertaining tour of the periodic table and the stories behind each element.",
            "For earth science, 'The Map That Changed the World' by Simon Winchester tells the story of William Smith and the birth of modern geology. 'The Sixth Extinction' by Elizabeth Kolbert examines the current biodiversity crisis through the lens of earth's previous mass extinction events.",
            "Online resources such as Khan Academy, MIT OpenCourseWare, and Coursera offer free courses covering all major scientific disciplines. The National Geographic Society, NASA, and the Smithsonian Institution provide extensive multimedia resources for learners of all ages.",
        ]),
        ("Appendix D: Methods and Instrumentation", [
            "The advancement of science has always been closely tied to the development of new instruments and methodologies. From Galileo's telescope to the Large Hadron Collider, each breakthrough in instrumentation has opened new domains of investigation.",
            "Microscopy techniques have evolved from simple optical microscopes capable of resolving structures down to about 200 nanometers to electron microscopes that can image individual atoms. Scanning tunneling microscopy and atomic force microscopy enable researchers to probe surfaces at the nanoscale, revealing details of crystal structure, molecular arrangement, and surface chemistry.",
            "Spectroscopic methods, including mass spectrometry, nuclear magnetic resonance, and infrared spectroscopy, allow scientists to determine the composition and structure of molecules with extraordinary precision. These techniques are indispensable in chemistry, biochemistry, and materials science.",
            "Computational methods have transformed every branch of science. Molecular dynamics simulations, climate models, and bioinformatics pipelines enable researchers to analyze datasets and test hypotheses that would be impossible to investigate through experimentation alone. The rise of machine learning and artificial intelligence promises to further accelerate scientific discovery.",
            "Remote sensing technologies, including satellites, LIDAR, and radar, provide global-scale data on atmospheric composition, land use, ocean temperature, and ice sheet dynamics. These tools are essential for monitoring environmental change and informing policy decisions.",
        ]),
    ]

    for sec_title, paras in additional_sections:
        doc.add_heading(sec_title, level=1)
        for para_text in paras:
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = 'Liberation Serif'
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = 1.15
        doc.add_page_break()

    # NO INDEX generated - that's the agent's task

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
