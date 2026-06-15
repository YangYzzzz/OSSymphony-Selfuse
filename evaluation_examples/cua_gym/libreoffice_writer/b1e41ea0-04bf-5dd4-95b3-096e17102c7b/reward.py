"""
Reward Script: Create a custom alphabetical index with letter headings and dot leaders
Task ID: writer_af_038
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): "Alphabetical Index" heading exists (Heading 1 style)
  Component 2 (0.25): Bold letter headings (single capital letters) for at least 5 distinct letters
  Component 3 (0.25): Index entries use RIGHT-aligned tab stop with DOTS leader
  Component 4 (0.25): Entries are indented below letter headings and follow term + page format
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_af_038'


def persist_app_state(domain):
    """Save any unsaved GUI state before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the "Alphabetical Index" heading
    index_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if 'alphabetical index' in p.text.strip().lower() and p.style and 'Heading' in p.style.name:
            index_heading_idx = i
            break

    # Component 1: "Alphabetical Index" heading exists with Heading style (0.25 points)
    try:
        if index_heading_idx is not None:
            heading_para = doc.paragraphs[index_heading_idx]
            print(f"PASS: Component 1 -- 'Alphabetical Index' heading found at paragraph {index_heading_idx}, style='{heading_para.style.name}' (0.25 pts)")
            total_score += 0.25
        else:
            print("FAIL: Component 1 -- No 'Alphabetical Index' heading with Heading style found")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    if index_heading_idx is None:
        # No index found at all, remaining components will all fail
        print("FAIL: Components 2-4 -- No index section to evaluate")
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Collect index paragraphs (everything after the heading)
    index_paras = doc.paragraphs[index_heading_idx + 1:]

    # Identify letter headings and entry paragraphs
    letter_headings = []  # list of single uppercase letters found as bold headings
    entry_paras = []  # paragraphs that look like index entries (term + tab + number)

    for p in index_paras:
        text = p.text.strip()
        if not text:
            continue

        # Check if it's a single uppercase letter (letter heading)
        if len(text) == 1 and text.isupper() and text.isalpha():
            # Check if bold
            has_bold = any(r.font.bold for r in p.runs if r.text.strip())
            letter_headings.append({'letter': text, 'bold': has_bold, 'para': p})
        elif '\t' in text:
            # Likely an index entry: "Term\tPageNum"
            entry_paras.append(p)

    # Component 2: Bold letter headings for at least 5 distinct letters (0.25 points)
    try:
        distinct_letters = set(lh['letter'] for lh in letter_headings)
        bold_letters = [lh for lh in letter_headings if lh['bold']]
        bold_count = len(bold_letters)

        if len(distinct_letters) >= 5 and bold_count >= 5:
            print(f"PASS: Component 2 -- {len(distinct_letters)} distinct letter headings found, {bold_count} are bold (0.25 pts)")
            total_score += 0.25
        elif len(distinct_letters) >= 3:
            # Partial credit: some letter headings present
            partial = 0.15
            print(f"PARTIAL: Component 2 -- {len(distinct_letters)} distinct letter headings (need >=5), {bold_count} bold ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 -- Only {len(distinct_letters)} letter headings found (need >=5), {bold_count} bold")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Index entries use RIGHT-aligned tab stop with DOTS leader (0.25 points)
    try:
        if len(entry_paras) == 0:
            print("FAIL: Component 3 -- No index entries found")
        else:
            entries_with_dot_leader = 0
            for ep in entry_paras:
                pf = ep.paragraph_format
                for ts in pf.tab_stops:
                    # Check for RIGHT alignment with DOTS leader
                    if ts.alignment == WD_TAB_ALIGNMENT.RIGHT and ts.leader == WD_TAB_LEADER.DOTS:
                        entries_with_dot_leader += 1
                        break

            ratio = entries_with_dot_leader / len(entry_paras)
            if ratio >= 0.8:
                print(f"PASS: Component 3 -- {entries_with_dot_leader}/{len(entry_paras)} entries have RIGHT tab with DOTS leader (0.25 pts)")
                total_score += 0.25
            elif ratio >= 0.4:
                partial = 0.15
                print(f"PARTIAL: Component 3 -- {entries_with_dot_leader}/{len(entry_paras)} entries have dot leaders ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 3 -- Only {entries_with_dot_leader}/{len(entry_paras)} entries have dot leaders")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Entries are indented and follow "Term\tPageNumber" format (0.25 points)
    try:
        if len(entry_paras) == 0:
            print("FAIL: Component 4 -- No index entries found")
        else:
            indented_count = 0
            valid_format_count = 0
            for ep in entry_paras:
                text = ep.text.strip()
                # Check indentation
                left_indent = ep.paragraph_format.left_indent
                if left_indent is not None and left_indent > 0:
                    indented_count += 1

                # Check format: term (text) + tab + page number (digits)
                parts = text.split('\t')
                if len(parts) == 2 and parts[0].strip() and parts[1].strip().isdigit():
                    valid_format_count += 1

            indent_ratio = indented_count / len(entry_paras)
            format_ratio = valid_format_count / len(entry_paras)

            if indent_ratio >= 0.8 and format_ratio >= 0.8:
                print(f"PASS: Component 4 -- {indented_count}/{len(entry_paras)} indented, {valid_format_count}/{len(entry_paras)} valid format (0.25 pts)")
                total_score += 0.25
            elif indent_ratio >= 0.5 or format_ratio >= 0.5:
                partial = 0.15
                print(f"PARTIAL: Component 4 -- {indented_count}/{len(entry_paras)} indented, {valid_format_count}/{len(entry_paras)} valid format ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 4 -- {indented_count}/{len(entry_paras)} indented, {valid_format_count}/{len(entry_paras)} valid format")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
