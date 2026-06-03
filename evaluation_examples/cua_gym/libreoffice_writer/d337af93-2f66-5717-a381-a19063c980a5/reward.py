"""
Reward Script: Add a page break before the section titled 'EXHIBIT A'
Task ID: writer_legal_011
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6 pts): Page break detected before EXHIBIT A paragraph (any method)
  Component 2 (0.4 pts): EXHIBIT A heading text preserved AND page break is correctly placed
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_011'


def persist_app_state(domain: str):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def find_exhibit_a_paragraph(doc):
    """Find the paragraph index containing 'EXHIBIT A' heading."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == 'EXHIBIT A':
            return i
    # Fallback: look for paragraph starting with EXHIBIT A
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('EXHIBIT A'):
            return i
    return -1


def has_page_break_before_exhibit_a(doc, exhibit_idx):
    """
    Check if there is a page break before the EXHIBIT A paragraph.
    Accepts three methods:
      1. page_break_before property on the EXHIBIT A paragraph
      2. A run-level <w:br w:type="page"/> in the last run of the preceding paragraph
      3. A run-level <w:br w:type="page"/> in a run of the EXHIBIT A paragraph itself (before text)
      4. A section break (new page) before EXHIBIT A
    Returns: (bool, str) -- whether break found and which method
    """
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    exhibit_para = doc.paragraphs[exhibit_idx]

    # Method 1: page_break_before paragraph property
    if exhibit_para.paragraph_format.page_break_before:
        return True, "page_break_before property on EXHIBIT A paragraph"

    # Method 2: Check XML for pageBreakBefore element directly
    pPr = exhibit_para._element.find('.//w:pPr', ns)
    if pPr is not None:
        pbb_elem = pPr.find('w:pageBreakBefore', ns)
        if pbb_elem is not None:
            # Check it's not explicitly set to false
            val = pbb_elem.get(f'{{{ns["w"]}}}val')
            if val is None or val.lower() in ('true', '1', 'on'):
                return True, "pageBreakBefore XML element on EXHIBIT A paragraph"

    # Method 3: Run-level page break in EXHIBIT A paragraph
    for run in exhibit_para.runs:
        for br in run.element.findall('.//w:br', ns):
            if br.attrib.get(f'{{{ns["w"]}}}type') == 'page':
                return True, "run-level page break in EXHIBIT A paragraph"

    # Method 4: Run-level page break at end of preceding paragraph
    if exhibit_idx > 0:
        prev_para = doc.paragraphs[exhibit_idx - 1]
        for run in prev_para.runs:
            for br in run.element.findall('.//w:br', ns):
                if br.attrib.get(f'{{{ns["w"]}}}type') == 'page':
                    return True, "run-level page break at end of preceding paragraph"

    # Method 5: Section break before EXHIBIT A (new page type)
    # Check if EXHIBIT A paragraph has a section break in its pPr
    if pPr is not None:
        sectPr = pPr.find('w:sectPr', ns)
        if sectPr is not None:
            sect_type = sectPr.find('w:type', ns)
            if sect_type is not None:
                val = sect_type.get(f'{{{ns["w"]}}}val')
                if val in ('nextPage', 'oddPage', 'evenPage'):
                    return True, f"section break ({val}) before EXHIBIT A"

    # Method 6: Check preceding paragraph for section break
    if exhibit_idx > 0:
        prev_pPr = doc.paragraphs[exhibit_idx - 1]._element.find('.//w:pPr', ns)
        if prev_pPr is not None:
            sectPr = prev_pPr.find('w:sectPr', ns)
            if sectPr is not None:
                sect_type = sectPr.find('w:type', ns)
                if sect_type is not None:
                    val = sect_type.get(f'{{{ns["w"]}}}val')
                    if val in ('nextPage', 'oddPage', 'evenPage'):
                        return True, f"section break ({val}) in preceding paragraph"
                else:
                    # Default section break type is nextPage
                    return True, "section break (default nextPage) in preceding paragraph"

    return False, "no page break found"


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the EXHIBIT A paragraph
    exhibit_idx = find_exhibit_a_paragraph(doc)
    if exhibit_idx < 0:
        print("CRITICAL: Could not find 'EXHIBIT A' paragraph in document")
        print("REWARD: 0.0")
        return 0.0

    print(f"INFO: Found 'EXHIBIT A' at paragraph index {exhibit_idx}")

    # Component 1: Page break detected before EXHIBIT A (0.6 points)
    # This is the core task requirement. Accepts any valid page break method.
    try:
        has_break, method = has_page_break_before_exhibit_a(doc, exhibit_idx)
        if has_break:
            print(f"PASS: Component 1 -- Page break found before EXHIBIT A via: {method} (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 -- No page break found before EXHIBIT A ({method})")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: EXHIBIT A heading preserved AND page break correctly placed (0.4 points)
    # Verifies that the page break insertion did not corrupt the EXHIBIT A heading
    # and that subsequent content is still intact.
    try:
        exhibit_para = doc.paragraphs[exhibit_idx]
        exhibit_text = exhibit_para.text.strip()

        # Sub-check 2a: EXHIBIT A text is preserved
        text_ok = (exhibit_text == 'EXHIBIT A')

        # Sub-check 2b: Content after EXHIBIT A is intact (at least the SOW title line)
        next_para_ok = (
            exhibit_idx + 1 < len(doc.paragraphs)
            and any(
                kw in doc.paragraphs[exhibit_idx + 1].text.strip().upper()
                for kw in ('STATEMENT OF WORK', 'MILESTONE')
            )
        )

        # This component only awards points if the page break IS present (anchored to the change)
        if has_break and text_ok and next_para_ok:
            print(f"PASS: Component 2 -- EXHIBIT A text preserved ('{exhibit_text}'), "
                  f"subsequent content intact, and page break correctly placed (0.4 pts)")
            total_score += 0.4
        elif not has_break:
            print(f"FAIL: Component 2 -- No page break present (prerequisite for this component)")
        elif not text_ok:
            print(f"FAIL: Component 2 -- EXHIBIT A text altered: '{exhibit_text}'")
        elif not next_para_ok:
            next_text_display = doc.paragraphs[exhibit_idx + 1].text.strip()[:60] if exhibit_idx + 1 < len(doc.paragraphs) else '(none)'
            print(f"FAIL: Component 2 -- Content after EXHIBIT A may be corrupted: '{next_text_display}'")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
