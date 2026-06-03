"""
Initial Setup: Add a page break before EXHIBIT A section
Task ID: writer_legal_011
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_011'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- TITLE ---
    title = doc.add_heading('PROFESSIONAL SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Parties ---
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Between')
    run.font.size = Pt(11)
    run.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run('Meridian Technology Solutions, Inc.')
    run2.bold = True
    run2.font.size = Pt(13)

    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = p3.add_run('and')
    run3.font.size = Pt(11)
    run3.italic = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = p4.add_run('Westbrook Financial Group, LLC')
    run4.bold = True
    run4.font.size = Pt(13)

    p5 = doc.add_paragraph()
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run5 = p5.add_run('Effective Date: January 15, 2025')
    run5.font.size = Pt(11)

    doc.add_paragraph('')  # spacer

    # --- RECITALS ---
    doc.add_heading('RECITALS', level=1)

    doc.add_paragraph(
        'WHEREAS, Meridian Technology Solutions, Inc. ("Provider") is a corporation '
        'organized and existing under the laws of the State of Delaware, with its '
        'principal place of business at 4200 Innovation Drive, Suite 800, San Jose, '
        'California 95134, and is engaged in the business of providing enterprise '
        'software development, cloud infrastructure management, and technical '
        'consulting services;'
    )

    doc.add_paragraph(
        'WHEREAS, Westbrook Financial Group, LLC ("Client") is a limited liability '
        'company organized and existing under the laws of the State of New York, with '
        'its principal place of business at 280 Park Avenue, 38th Floor, New York, '
        'New York 10017, and desires to engage Provider for certain professional '
        'services as described herein;'
    )

    doc.add_paragraph(
        'WHEREAS, Provider possesses the requisite expertise, personnel, and resources '
        'to perform the services contemplated under this Agreement, and Client wishes '
        'to retain Provider on the terms and conditions set forth below;'
    )

    doc.add_paragraph(
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements '
        'hereinafter set forth, and for other good and valuable consideration, the '
        'receipt and sufficiency of which are hereby acknowledged, the parties agree '
        'as follows:'
    )

    # --- ARTICLE 1: SCOPE OF SERVICES ---
    doc.add_heading('ARTICLE 1: SCOPE OF SERVICES', level=1)

    doc.add_paragraph(
        '1.1  Provider shall furnish the professional services described in Exhibit A '
        'attached hereto and incorporated herein by reference (the "Services"). The '
        'Services shall include, but not be limited to, custom software development, '
        'system integration, data migration, quality assurance testing, and ongoing '
        'technical support as further detailed in the Statement of Work.'
    )

    doc.add_paragraph(
        '1.2  Provider shall perform the Services in a professional and workmanlike '
        'manner, consistent with generally accepted industry standards and practices. '
        'Provider shall assign qualified personnel with appropriate experience and '
        'expertise to perform the Services.'
    )

    doc.add_paragraph(
        '1.3  Any changes to the scope of Services shall be documented in a written '
        'Change Order signed by authorized representatives of both parties. No '
        'additional compensation shall be due for work performed outside the approved '
        'scope unless a Change Order has been duly executed.'
    )

    # --- ARTICLE 2: COMPENSATION ---
    doc.add_heading('ARTICLE 2: COMPENSATION AND PAYMENT', level=1)

    doc.add_paragraph(
        '2.1  In consideration of the Services rendered, Client shall pay Provider a '
        'total fee of Three Hundred Seventy-Five Thousand Dollars ($375,000.00) '
        '(the "Fee"), payable in accordance with the milestone schedule set forth in '
        'Exhibit A.'
    )

    doc.add_paragraph(
        '2.2  Provider shall submit monthly invoices to Client for Services performed '
        'during the preceding calendar month. Each invoice shall include a detailed '
        'description of the Services performed, the hours expended by each assigned '
        'team member, and any pre-approved reimbursable expenses.'
    )

    doc.add_paragraph(
        '2.3  Client shall remit payment within thirty (30) calendar days of receipt '
        'of each properly submitted invoice. Late payments shall accrue interest at '
        'the rate of one and one-half percent (1.5%) per month, or the maximum rate '
        'permitted by applicable law, whichever is less.'
    )

    # --- ARTICLE 3: TERM AND TERMINATION ---
    doc.add_heading('ARTICLE 3: TERM AND TERMINATION', level=1)

    doc.add_paragraph(
        '3.1  This Agreement shall commence on the Effective Date and shall continue '
        'for a period of twenty-four (24) months, unless earlier terminated in '
        'accordance with the provisions of this Article (the "Term").'
    )

    doc.add_paragraph(
        '3.2  Either party may terminate this Agreement for cause upon sixty (60) days '
        'written notice to the other party in the event of a material breach of this '
        'Agreement, provided that the breaching party has failed to cure such breach '
        'within the sixty-day notice period.'
    )

    doc.add_paragraph(
        '3.3  Client may terminate this Agreement for convenience upon ninety (90) '
        'days prior written notice to Provider. In the event of such termination, '
        'Client shall pay Provider for all Services satisfactorily performed through '
        'the effective date of termination, plus any reasonable wind-down costs.'
    )

    # --- ARTICLE 4: CONFIDENTIALITY ---
    doc.add_heading('ARTICLE 4: CONFIDENTIALITY', level=1)

    doc.add_paragraph(
        '4.1  Each party acknowledges that during the performance of this Agreement, '
        'it may receive or have access to confidential and proprietary information of '
        'the other party ("Confidential Information"). Confidential Information '
        'includes, without limitation, trade secrets, business plans, financial data, '
        'customer lists, technical specifications, and software source code.'
    )

    doc.add_paragraph(
        '4.2  The receiving party shall hold all Confidential Information in strict '
        'confidence and shall not disclose such information to any third party without '
        'the prior written consent of the disclosing party. The receiving party shall '
        'use the Confidential Information solely for the purpose of performing its '
        'obligations under this Agreement.'
    )

    # --- ARTICLE 5: SIGNATURES ---
    doc.add_heading('ARTICLE 5: GENERAL PROVISIONS', level=1)

    doc.add_paragraph(
        '5.1  This Agreement constitutes the entire agreement between the parties with '
        'respect to the subject matter hereof and supersedes all prior negotiations, '
        'representations, warranties, commitments, offers, contracts, and writings, '
        'whether written or oral, with respect to the subject matter hereof.'
    )

    doc.add_paragraph(
        '5.2  This Agreement shall be governed by and construed in accordance with the '
        'laws of the State of New York, without regard to its conflict of laws '
        'principles. Any disputes arising under this Agreement shall be resolved in '
        'the state or federal courts located in the Borough of Manhattan, New York.'
    )

    # --- Signature Block ---
    doc.add_paragraph('')  # spacer

    sig_para = doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.')
    sig_para.paragraph_format.space_after = Pt(24)

    # Provider signature
    p_sig1 = doc.add_paragraph()
    run_sig1 = p_sig1.add_run('MERIDIAN TECHNOLOGY SOLUTIONS, INC.')
    run_sig1.bold = True
    doc.add_paragraph('By: ________________________________')
    doc.add_paragraph('Name: Jonathan R. Blackwell')
    doc.add_paragraph('Title: Chief Executive Officer')
    doc.add_paragraph('Date: January 15, 2025')

    doc.add_paragraph('')  # spacer

    # Client signature
    p_sig2 = doc.add_paragraph()
    run_sig2 = p_sig2.add_run('WESTBROOK FINANCIAL GROUP, LLC')
    run_sig2.bold = True
    doc.add_paragraph('By: ________________________________')
    doc.add_paragraph('Name: Catherine M. Thornton')
    doc.add_paragraph('Title: Managing Partner')
    doc.add_paragraph('Date: January 15, 2025')

    # --- EXHIBIT A (flows directly, NO page break) ---
    doc.add_paragraph('')  # spacer

    exhibit_heading = doc.add_heading('EXHIBIT A', level=1)
    exhibit_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # NO page_break_before set -- this is the initial state

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_sub = sub.add_run('STATEMENT OF WORK AND MILESTONE SCHEDULE')
    run_sub.bold = True
    run_sub.font.size = Pt(12)

    doc.add_paragraph('')

    doc.add_heading('1. Project Overview', level=2)
    doc.add_paragraph(
        'Provider shall design, develop, and deploy a comprehensive Enterprise '
        'Resource Planning (ERP) integration platform for Client. The platform shall '
        'connect Client\'s existing financial reporting systems with a new cloud-based '
        'data analytics dashboard, enabling real-time portfolio tracking and automated '
        'compliance reporting.'
    )

    doc.add_heading('2. Deliverables', level=2)
    doc.add_paragraph('Phase 1 - Requirements Analysis and System Architecture', style='List Bullet')
    doc.add_paragraph('Phase 2 - Core Platform Development and API Integration', style='List Bullet')
    doc.add_paragraph('Phase 3 - Data Migration and Quality Assurance Testing', style='List Bullet')
    doc.add_paragraph('Phase 4 - User Acceptance Testing and Deployment', style='List Bullet')
    doc.add_paragraph('Phase 5 - Post-Launch Support and Optimization', style='List Bullet')

    doc.add_heading('3. Milestone Schedule', level=2)

    # Milestone table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Milestone', 'Description', 'Due Date', 'Payment']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    milestones = [
        ['M1', 'Requirements Sign-Off', 'March 15, 2025', '$56,250'],
        ['M2', 'Architecture Approval', 'May 1, 2025', '$75,000'],
        ['M3', 'Beta Release', 'August 15, 2025', '$93,750'],
        ['M4', 'UAT Completion', 'October 31, 2025', '$93,750'],
        ['M5', 'Go-Live & Handover', 'December 15, 2025', '$56,250'],
    ]
    for r, row_data in enumerate(milestones, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph('')
    doc.add_paragraph(
        'Total Contract Value: $375,000.00. Payment terms are net-30 from milestone '
        'acceptance date. All amounts are in United States Dollars.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
