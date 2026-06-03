"""
Initial Setup: Legal clause document with 5 sections, third paragraph contains 10 clauses as one block
Task ID: osworld_writer_spacing_003
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'  # VM path — all scripts run on the VM
TASK_ID = 'osworld_writer_spacing_003'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Section 1: Introduction
    heading1 = doc.add_heading('SERVICE AGREEMENT', level=1)

    intro = doc.add_paragraph(
        'This Service Agreement ("Agreement") is entered into as of March 1, 2025, '
        'by and between Meridian Solutions Inc., a corporation organized under the laws '
        'of the State of Delaware ("Service Provider"), and Harrington & Associates LLC, '
        'a limited liability company ("Client"). Both parties agree to the following terms '
        'and conditions governing the provision of professional consulting services.'
    )

    # Section 2: Scope of Work
    doc.add_heading('SECTION 1: SCOPE OF WORK', level=2)

    scope = doc.add_paragraph(
        'The Service Provider agrees to deliver comprehensive IT infrastructure consulting '
        'services including system architecture review, security assessment, and performance '
        'optimization. All deliverables will be provided in written reports with supporting '
        'documentation. The Client shall provide access to all necessary systems, personnel, '
        'and data required for the Service Provider to fulfill its obligations under this Agreement.'
    )

    # Section 3: General Terms and Conditions — all 10 clauses in ONE paragraph (no empty lines)
    doc.add_heading('SECTION 2: GENERAL TERMS AND CONDITIONS', level=2)

    # 10 clauses as one single block paragraph — this is what the agent must split
    clauses_block = (
        'The contractor shall complete all work by the specified deadline. '
        'Payment will be made in 30-day installments upon receipt of invoice. '
        'Either party may terminate this agreement with 30 days written notice. '
        'All intellectual property developed under this agreement belongs to the client. '
        'The service provider shall maintain confidentiality of all proprietary information. '
        'Disputes shall be resolved through binding arbitration in the state of Delaware. '
        'The service provider shall carry professional liability insurance of no less than $2,000,000. '
        'Work product must meet the quality standards specified in Exhibit A attached hereto. '
        'The service provider may engage approved subcontractors with prior written consent. '
        'This agreement shall be governed by and construed in accordance with applicable law.'
    )
    clause_para = doc.add_paragraph(clauses_block)

    # Section 4: Payment Terms
    doc.add_heading('SECTION 3: PAYMENT TERMS', level=2)

    payment = doc.add_paragraph(
        'The Client agrees to pay the Service Provider a total contract value of $125,000 USD, '
        'payable in monthly installments of $12,500 over the ten-month engagement period. '
        'Invoices must be submitted by the 15th of each month and will be processed within '
        'thirty (30) business days. Late payments shall accrue interest at a rate of 1.5% '
        'per month. All fees are exclusive of applicable taxes, which shall be the sole '
        'responsibility of the Client.'
    )

    # Section 5: Limitation of Liability
    doc.add_heading('SECTION 4: LIMITATION OF LIABILITY', level=2)

    liability = doc.add_paragraph(
        'In no event shall either party be liable for indirect, incidental, special, or '
        'consequential damages arising out of or related to this Agreement, even if advised '
        'of the possibility of such damages. The total aggregate liability of either party '
        'shall not exceed the total fees paid under this Agreement during the twelve months '
        'preceding the claim. This limitation applies regardless of the cause of action, '
        'whether in contract, tort, warranty, or any other legal or equitable theory.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
