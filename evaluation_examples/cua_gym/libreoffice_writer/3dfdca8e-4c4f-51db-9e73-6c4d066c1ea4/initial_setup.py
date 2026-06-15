"""
Initial Setup: Meridian Health Case Study - plain unformatted draft
Task ID: writer_mktg_004
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_004'
OUTPUT = f'{WORKDIR}/meridian_case_study.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Default paragraph style: plain 12pt (no bold, no color, no special formatting)
    def add_plain(text, bold=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = False
        run.italic = False
        run.font.size = Pt(12)
        return para

    # Title — plain, not formatted (agent must format it)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Meridian Health Case Study")
    title_run.bold = False
    title_run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # Challenge section
    challenge_heading = doc.add_paragraph()
    ch_run = challenge_heading.add_run("Challenge")
    ch_run.bold = False
    ch_run.font.size = Pt(12)

    challenge_body = doc.add_paragraph()
    cb_run = challenge_body.add_run(
        "Meridian Health, a regional healthcare provider with over 120,000 active patients, "
        "struggled to drive meaningful adoption of their patient portal. Despite investing in "
        "digital infrastructure, portal engagement hovered below 30%, with patients defaulting "
        "to phone calls for appointments, prescription refills, and lab results. The clinical "
        "team spent an estimated 2,400 staff hours per month handling requests that the portal "
        "was designed to automate."
    )
    cb_run.bold = False
    cb_run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # Solution section
    solution_heading = doc.add_paragraph()
    sh_run = solution_heading.add_run("Solution")
    sh_run.bold = False
    sh_run.font.size = Pt(12)

    solution_body = doc.add_paragraph()
    sb_run = solution_body.add_run(
        "Meridian Health partnered with NovaTech Solutions to deploy a tailored patient "
        "engagement platform integrated directly with their existing EHR system. The rollout "
        "included personalized onboarding sequences, push notifications for upcoming appointments "
        "and medication reminders, and an intuitive mobile-first interface designed with input "
        "from Meridian's patient advisory board. NovaTech's implementation team provided "
        "dedicated training for clinical staff and customized the portal's messaging templates "
        "to match Meridian's brand voice."
    )
    sb_run.bold = False
    sb_run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # Results section
    results_heading = doc.add_paragraph()
    rh_run = results_heading.add_run("Results")
    rh_run.bold = False
    rh_run.font.size = Pt(12)

    results_body = doc.add_paragraph()
    rb_run = results_body.add_run(
        "Within the first quarter following deployment, Meridian Health recorded dramatic "
        "improvements across all key performance indicators. Patient portal adoption surged "
        "and the organization achieved measurable financial and operational gains."
    )
    rb_run.bold = False
    rb_run.font.size = Pt(12)

    # Metrics as inline text (plain, not bold, not 14pt) — agent must format them
    metrics_para = doc.add_paragraph()
    m1_run = metrics_para.add_run("47% engagement increase")
    m1_run.bold = False
    m1_run.font.size = Pt(12)
    metrics_para.add_run(" | ")
    m2_run = metrics_para.add_run("$2.1M annual savings")
    m2_run.bold = False
    m2_run.font.size = Pt(12)
    metrics_para.add_run(" | ")
    m3_run = metrics_para.add_run("89% patient satisfaction score")
    m3_run.bold = False
    m3_run.font.size = Pt(12)

    results_detail = doc.add_paragraph()
    rd_run = results_detail.add_run(
        "Staff hours spent on administrative phone requests dropped by 61%, freeing clinical "
        "personnel to focus on direct patient care. The portal's secure messaging feature "
        "processed over 18,000 non-urgent patient inquiries in the quarter, with an average "
        "response time of under four hours."
    )
    rd_run.bold = False
    rd_run.font.size = Pt(12)

    doc.add_paragraph()  # blank line

    # Client Testimonial section
    testimonial_heading = doc.add_paragraph()
    th_run = testimonial_heading.add_run("Client Testimonial")
    th_run.bold = False
    th_run.font.size = Pt(12)

    testimonial_para = doc.add_paragraph()
    tq_run = testimonial_para.add_run(
        "We saw a 47% increase in patient portal engagement within the first quarter. "
        "NovaTech's team understood our workflows and delivered a solution that our staff "
        "and patients actually use. The ROI has been remarkable and the impact on care "
        "coordination is something we didn't expect to see this quickly."
    )
    tq_run.bold = False
    tq_run.italic = False
    tq_run.font.size = Pt(12)

    attribution_para = doc.add_paragraph()
    attr_run = attribution_para.add_run(
        "— Dr. Sarah Lindstrom, Chief Medical Officer, Meridian Health"
    )
    attr_run.bold = False
    attr_run.italic = False
    attr_run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
