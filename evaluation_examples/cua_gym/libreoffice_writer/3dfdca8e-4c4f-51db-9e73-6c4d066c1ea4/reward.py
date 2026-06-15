"""
Reward Script: writer_mktg_004 — Meridian Health Case Study Formatting
Task ID: writer_mktg_004
Domain: libreoffice_writer
Scoring:
  Component 1: Title formatted as 22pt bold, centered                        — 0.20 pts
  Component 2: Section headers (Challenge/Solution/Results) 16pt bold #0A2463 — 0.25 pts
  Component 3: Key metrics paragraph (3 metrics) bold at 14pt               — 0.20 pts
  Component 4: Pull quote in 1x1 table with F2F2F2 shading, blue left border, italic text — 0.25 pts
  Component 5: Page header "NovaTech Solutions" right-aligned, ~10pt gray    — 0.10 pts
  Total: 1.00
"""

import os
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_004'
FILE_PATH = f'{WORKDIR}/meridian_case_study.docx'


def color_distance(c1_hex, r2, g2, b2):
    """Euclidean distance between an RGB hex string and explicit RGB values."""
    r1 = int(c1_hex[0:2], 16)
    g1 = int(c1_hex[2:4], 16)
    b1 = int(c1_hex[4:6], 16)
    return ((r1 - r2) ** 2 + (g1 - g2) ** 2 + (b1 - b2) ** 2) ** 0.5


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Title is 22pt bold and centered (0.20 points)
    # The title "Meridian Health Case Study" is Para 0.
    # Initial: 12pt, not bold, not centered. Golden: 22pt, bold, centered.
    # -------------------------------------------------------------------------
    try:
        title_para = None
        for para in doc.paragraphs:
            if 'Meridian Health Case Study' in para.text:
                title_para = para
                break

        if title_para is None:
            print("FAIL: Component 1 — Title paragraph 'Meridian Health Case Study' not found")
        else:
            title_runs = [r for r in title_para.runs if r.text.strip()]
            if not title_runs:
                print("FAIL: Component 1 — Title paragraph has no runs with text")
            else:
                # Check bold: all substantive runs must be bold
                all_bold = all(r.bold is True for r in title_runs)
                # Check size: all substantive runs must be 22pt
                all_22pt = all(
                    r.font.size is not None and abs(r.font.size.pt - 22.0) < 0.5
                    for r in title_runs
                )
                # Check alignment: centered
                alignment_ok = (
                    title_para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                )

                if all_bold and all_22pt and alignment_ok:
                    print(f"PASS: Component 1 — Title is 22pt bold and centered (0.20 pts)")
                    total_score += 0.20
                else:
                    reasons = []
                    if not all_bold:
                        bold_vals = [r.bold for r in title_runs]
                        reasons.append(f"bold={bold_vals}")
                    if not all_22pt:
                        sizes = [r.font.size.pt if r.font.size else None for r in title_runs]
                        reasons.append(f"sizes={sizes}")
                    if not alignment_ok:
                        reasons.append(f"alignment={title_para.paragraph_format.alignment}")
                    print(f"FAIL: Component 1 — Title formatting issues: {', '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Section headers (Challenge, Solution, Results) are 16pt bold
    #              with color #0A2463 (0.25 points)
    # Initial: 12pt, not bold, no color. Golden: 16pt, bold, #0A2463.
    # -------------------------------------------------------------------------
    try:
        SECTION_HEADERS = ['Challenge', 'Solution', 'Results']
        TARGET_COLOR_HEX = '0A2463'
        TARGET_R, TARGET_G, TARGET_B = 0x0A, 0x24, 0x63

        headers_found = {h: False for h in SECTION_HEADERS}
        headers_ok = {h: False for h in SECTION_HEADERS}

        for para in doc.paragraphs:
            text = para.text.strip()
            if text in SECTION_HEADERS:
                headers_found[text] = True
                runs = [r for r in para.runs if r.text.strip()]
                if not runs:
                    print(f"  FAIL: Header '{text}' has no styled runs")
                    continue

                is_bold = all(r.bold is True for r in runs)
                is_16pt = all(
                    r.font.size is not None and abs(r.font.size.pt - 16.0) < 0.5
                    for r in runs
                )
                # Color check: accept approximate match
                has_blue_color = False
                for r in runs:
                    try:
                        if r.font.color and r.font.color.type:
                            c = str(r.font.color.rgb)  # hex string like '0A2463'
                            dist = color_distance(c, TARGET_R, TARGET_G, TARGET_B)
                            if dist < 30:
                                has_blue_color = True
                                break
                    except Exception:
                        pass

                if is_bold and is_16pt and has_blue_color:
                    headers_ok[text] = True
                else:
                    reasons = []
                    if not is_bold:
                        reasons.append(f"bold={[r.bold for r in runs]}")
                    if not is_16pt:
                        sizes = [r.font.size.pt if r.font.size else None for r in runs]
                        reasons.append(f"size={sizes}")
                    if not has_blue_color:
                        colors = []
                        for r in runs:
                            try:
                                if r.font.color and r.font.color.type:
                                    colors.append(str(r.font.color.rgb))
                                else:
                                    colors.append(None)
                            except Exception:
                                colors.append('err')
                        reasons.append(f"color={colors}")
                    print(f"  FAIL: Header '{text}' — {', '.join(reasons)}")

        num_ok = sum(headers_ok.values())
        num_found = sum(headers_found.values())

        if num_ok == 3:
            print(f"PASS: Component 2 — All 3 section headers are 16pt bold #0A2463 (0.25 pts)")
            total_score += 0.25
        elif num_ok == 2:
            print(f"PARTIAL: Component 2 — {num_ok}/3 section headers formatted correctly (0.15 pts)")
            total_score += 0.15
        elif num_ok == 1:
            print(f"PARTIAL: Component 2 — {num_ok}/3 section headers formatted correctly (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 2 — 0/3 section headers formatted correctly (found={num_found})")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Key metrics paragraph has all three metrics bold at 14pt (0.20 pts)
    # Para 10 contains: "47% engagement increase | $2.1M annual savings | 89% patient satisfaction score"
    # Initial: 12pt, not bold. Golden: 14pt, bold for all metric runs.
    # -------------------------------------------------------------------------
    try:
        METRICS_KEYWORDS = ['47% engagement', '$2.1M annual', '89% patient satisfaction']

        metrics_para = None
        for para in doc.paragraphs:
            if '47% engagement' in para.text and '$2.1M' in para.text and '89% patient' in para.text:
                metrics_para = para
                break

        if metrics_para is None:
            print("FAIL: Component 3 — Metrics paragraph not found")
        else:
            # Find runs that contain metric text (not separator ' | ')
            metric_runs = []
            for r in metrics_para.runs:
                if r.text.strip() and '|' not in r.text:
                    metric_runs.append(r)

            if not metric_runs:
                print("FAIL: Component 3 — No metric runs found in paragraph")
            else:
                all_bold = all(r.bold is True for r in metric_runs)
                all_14pt = all(
                    r.font.size is not None and abs(r.font.size.pt - 14.0) < 0.5
                    for r in metric_runs
                )

                if all_bold and all_14pt:
                    print(f"PASS: Component 3 — All {len(metric_runs)} metric runs are 14pt bold (0.20 pts)")
                    total_score += 0.20
                else:
                    reasons = []
                    if not all_bold:
                        bold_vals = [(r.text[:20], r.bold) for r in metric_runs]
                        reasons.append(f"bold={bold_vals}")
                    if not all_14pt:
                        size_vals = [(r.text[:20], r.font.size.pt if r.font.size else None) for r in metric_runs]
                        reasons.append(f"size={size_vals}")
                    print(f"FAIL: Component 3 — Metric runs not formatted correctly: {', '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: Pull quote (Dr. Lindstrom quote) in a 1x1 table with
    #              F2F2F2 shading, blue (#0A2463) left border, italic text (0.25 pts)
    # Initial: Plain paragraph at para 14. Golden: inside a 1x1 table.
    # -------------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 4 — No tables found (expected 1 table with pull quote)")
        else:
            table = doc.tables[0]
            if len(table.rows) < 1 or len(table.columns) < 1:
                print("FAIL: Component 4 — Table has unexpected shape")
            else:
                cell = table.cell(0, 0)
                cell_text = cell.text
                has_lindstrom_quote = 'Lindstrom' in cell_text or '47% increase in patient portal' in cell_text

                if not has_lindstrom_quote:
                    print(f"FAIL: Component 4 — Table does not contain Dr. Lindstrom quote (cell text: {repr(cell_text[:60])})")
                else:
                    # Check italic formatting of quote text
                    all_italic = True
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                if run.italic is not True:
                                    all_italic = False
                                    break

                    # Check F2F2F2 shading
                    tc = cell._tc
                    tc_xml = tc.xml
                    has_shading = False
                    if 'w:shd' in tc_xml:
                        # Look for fill=F2F2F2 (case insensitive)
                        shd_match = re.findall(r'w:fill="([^"]+)"', tc_xml)
                        for fill_val in shd_match:
                            if fill_val.upper() in ('F2F2F2', 'F2F2F2F2'):
                                has_shading = True
                                break

                    # Check left border color #0A2463
                    has_blue_left_border = False
                    if 'w:tcBorders' in tc_xml:
                        left_match = re.findall(r'<w:left[^/]*/>', tc_xml)
                        for lm in left_match:
                            color_match = re.search(r'w:color="([^"]+)"', lm)
                            if color_match:
                                border_color = color_match.group(1).upper()
                                if border_color in ('0A2463', '0a2463'):
                                    has_blue_left_border = True
                                    break
                            if 'w:val=' in lm and 'single' in lm:
                                # Still check color
                                pass

                    sub_score = 0.0
                    if all_italic:
                        sub_score += 0.10
                        print(f"  PASS: Pull quote text is italic (+0.10)")
                    else:
                        print(f"  FAIL: Pull quote text is not fully italic")

                    if has_shading:
                        sub_score += 0.10
                        print(f"  PASS: Table cell has F2F2F2 shading (+0.10)")
                    else:
                        fill_vals_found = re.findall(r'w:fill="([^"]+)"', tc_xml)
                        print(f"  FAIL: Table cell missing F2F2F2 shading (found fills: {fill_vals_found})")

                    if has_blue_left_border:
                        sub_score += 0.05
                        print(f"  PASS: Table cell left border is #0A2463 blue (+0.05)")
                    else:
                        print(f"  FAIL: Table cell left border not #0A2463 blue")

                    if sub_score >= 0.25:
                        sub_score = 0.25
                    if sub_score > 0:
                        print(f"Component 4 — Pull quote sub-score: {sub_score}/0.25")
                        total_score += sub_score
                    else:
                        print("FAIL: Component 4 — Pull quote not properly formatted")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Page header "NovaTech Solutions" right-aligned, ~10pt, gray (0.10 pts)
    # Initial: Header is linked/empty. Golden: Header contains "NovaTech Solutions" right-aligned.
    # -------------------------------------------------------------------------
    try:
        section = doc.sections[0]
        header = section.header
        header_text = ''
        header_align = None
        header_size = None
        header_color_hex = None

        for para in header.paragraphs:
            if para.text.strip():
                header_text = para.text.strip()
                header_align = para.paragraph_format.alignment
                for run in para.runs:
                    if run.text.strip():
                        header_size = run.font.size.pt if run.font.size else None
                        try:
                            if run.font.color and run.font.color.type:
                                header_color_hex = str(run.font.color.rgb)
                        except Exception:
                            pass
                break

        has_novatech = 'NovaTech Solutions' in header_text or 'novatech solutions' in header_text.lower()
        is_right_aligned = (header_align == WD_PARAGRAPH_ALIGNMENT.RIGHT)
        is_approx_10pt = (header_size is not None and abs(header_size - 10.0) < 1.5)

        # Gray: #808080 or similar (R==G==B, relatively dark)
        is_gray = False
        if header_color_hex:
            try:
                r = int(header_color_hex[0:2], 16)
                g = int(header_color_hex[2:4], 16)
                b = int(header_color_hex[4:6], 16)
                # Gray: all channels similar and in range 80-160
                if abs(r - g) < 30 and abs(g - b) < 30 and 50 < r < 200:
                    is_gray = True
            except Exception:
                pass

        if has_novatech and is_right_aligned:
            print(f"PASS: Component 5 — Header contains 'NovaTech Solutions' right-aligned (0.10 pts)")
            print(f"  Header text={repr(header_text)}, size={header_size}pt, color={header_color_hex}, gray={is_gray}")
            total_score += 0.10
        else:
            reasons = []
            if not has_novatech:
                reasons.append(f"header text is '{header_text}' (expected 'NovaTech Solutions')")
            if not is_right_aligned:
                reasons.append(f"alignment={header_align} (expected RIGHT)")
            print(f"FAIL: Component 5 — Header issues: {', '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -------------------------------------------------------------------------
    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score:.4f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
