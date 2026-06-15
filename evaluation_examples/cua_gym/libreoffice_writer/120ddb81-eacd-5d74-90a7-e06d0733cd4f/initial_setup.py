"""
Initial Setup: Federal court filing caption block - plain text only
Task ID: writer_legal_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_052'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # All content as plain, unformatted text - no bold, no centering, no tables
    doc.add_paragraph('UNITED STATES DISTRICT COURT NORTHERN DISTRICT OF CALIFORNIA')
    doc.add_paragraph('')  # blank line
    doc.add_paragraph('JOHN DOE, Plaintiff,')
    doc.add_paragraph('v.')
    doc.add_paragraph('ACME CORPORATION, Defendant.')
    doc.add_paragraph('')  # blank line
    doc.add_paragraph('Case No. 3:24-cv-01234-ABC')
    doc.add_paragraph('Hon. Jane Smith')
    doc.add_paragraph('MOTION FOR SUMMARY JUDGMENT')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
