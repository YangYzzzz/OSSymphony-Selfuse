"""
Reward Script: Format federal court filing caption block
Task ID: writer_legal_052
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Court name centered and bold
  Component 2 (0.25): 2-column, 1-row table exists
  Component 3 (0.25): Correct content in table cells with 'v.' centered
  Component 4 (0.25): Border config - outer borders hidden, vertical divider present
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_052'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Court name is centered and bold (0.25 points)
    # The golden state has the court name split across two paragraphs,
    # both centered and bold. The initial state has it as a single
    # unformatted paragraph with no centering and no bold.
    try:
        # Find paragraphs containing court name text (before the table)
        court_paras = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and ('UNITED STATES DISTRICT COURT' in text or
                         'NORTHERN DISTRICT OF CALIFORNIA' in text):
                court_paras.append(para)

        if len(court_paras) >= 1:
            centered_count = sum(1 for cp in court_paras
                                 if cp.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
            bold_count = sum(1 for cp in court_paras
                            if all(r.font.bold for r in cp.runs if r.text.strip()))
            all_centered = (centered_count == len(court_paras))  # derived from check
            all_bold = (bold_count == len(court_paras))  # derived from check

            if all_centered and all_bold:
                print(f"PASS: Component 1 — Court name centered and bold ({len(court_paras)} paras) (0.25 pts)")
                total_score += 0.25
            elif all_centered:
                print(f"PARTIAL: Component 1 — Court name centered but not bold (0.1 pts)")
                total_score += 0.1
            elif all_bold:
                print(f"PARTIAL: Component 1 — Court name bold but not centered (0.1 pts)")
                total_score += 0.1
            else:
                print(f"FAIL: Component 1 — Court name not centered (centered={all_centered}) and not bold (bold={all_bold})")
        else:
            print("FAIL: Component 1 — Court name paragraphs not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Table with 2 columns and 1 row exists (0.25 points)
    # Initial state has 0 tables. Golden state has exactly 1 table with 2 cols, 1 row.
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            num_rows = len(table.rows)
            num_cols = len(table.columns)
            if num_cols == 2 and num_rows >= 1:
                print(f"PASS: Component 2 — Table found with {num_rows} row(s) and {num_cols} columns (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 2 — Table has {num_rows} rows, {num_cols} cols (expected 1 row, 2 cols)")
        else:
            print(f"FAIL: Component 2 — No tables found (expected 1)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Correct content in table cells (0.25 points)
    # Left cell: party names with 'v.' centered
    # Right cell: case number, judge, motion title
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)

            left_text = left_cell.text.strip()
            right_text = right_cell.text.strip()

            sub_score = 0.0

            # Left cell should contain party names and "v."
            has_plaintiff = 'JOHN DOE' in left_text or 'Plaintiff' in left_text
            has_defendant = 'ACME CORPORATION' in left_text or 'Defendant' in left_text
            has_v = 'v.' in left_text

            if has_plaintiff and has_defendant and has_v:
                # Check if 'v.' paragraph is centered
                v_para_centered = any(
                    para.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                    for para in left_cell.paragraphs
                    if para.text.strip() == 'v.'
                )

                if v_para_centered:
                    sub_score += 0.15
                    print(f"PASS: Component 3a — Left cell has parties with 'v.' centered")
                elif not v_para_centered:
                    sub_score += 0.05
                    print(f"PARTIAL: Component 3a — Left cell has parties but 'v.' not centered")
            else:
                print(f"FAIL: Component 3a — Left cell missing party info (plaintiff={has_plaintiff}, defendant={has_defendant}, v={has_v})")

            # Right cell should contain case info
            has_case_no = '3:24-cv-01234-ABC' in right_text
            has_judge = 'Jane Smith' in right_text
            has_motion = 'MOTION FOR SUMMARY JUDGMENT' in right_text

            if has_case_no and has_judge and has_motion:
                sub_score += 0.10
                print(f"PASS: Component 3b — Right cell has case number, judge, and motion title")
            else:
                print(f"FAIL: Component 3b — Right cell missing info (case={has_case_no}, judge={has_judge}, motion={has_motion})")

            if sub_score > 0:
                total_score += sub_score
            print(f"Component 3 total: {sub_score}/0.25 pts")
        else:
            print("FAIL: Component 3 — No table to check cell content")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Border configuration (0.25 points)
    # Outer borders should be hidden (nil), vertical divider should be present (single)
    # Initial state has no table at all, so this will fail on initial.
    try:
        if len(doc.tables) >= 1:
            table = doc.tables[0]
            tbl = table._tbl
            tblPr = tbl.find(qn('w:tblPr'))

            sub_score = 0.0

            if tblPr is not None:
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is not None:
                    # Check vertical divider exists
                    insideV = tblBorders.find(qn('w:insideV'))
                    if insideV is not None and insideV.get(qn('w:val')) == 'single':
                        sub_score += 0.10
                        print("PASS: Component 4a — Vertical divider is present (single)")
                    else:
                        # Also check cell-level borders for vertical divider
                        cell_right = table.cell(0, 0)._tc.find(qn('w:tcPr'))
                        if cell_right is not None:
                            tcBorders = cell_right.find(qn('w:tcBorders'))
                            if tcBorders is not None:
                                right_border = tcBorders.find(qn('w:right'))
                                if right_border is not None and right_border.get(qn('w:val')) == 'single':
                                    sub_score += 0.10
                                    print("PASS: Component 4a — Vertical divider found on cell border (single)")
                                else:
                                    print(f"FAIL: Component 4a — No vertical divider found")
                            else:
                                print("FAIL: Component 4a — No cell borders found")
                        else:
                            print("FAIL: Component 4a — No vertical divider found")

                    # Check outer borders are hidden (nil or none)
                    visible_outer = [
                        border_name for border_name in ['top', 'left', 'bottom', 'right']
                        if tblBorders.find(qn(f'w:{border_name}')) is not None
                        and tblBorders.find(qn(f'w:{border_name}')).get(qn('w:val')) not in ('nil', 'none')
                    ]

                    if len(visible_outer) == 0:
                        sub_score += 0.15
                        print("PASS: Component 4b — Outer borders are hidden")
                    else:
                        print("FAIL: Component 4b — Some outer borders are not hidden")
                else:
                    # No explicit table borders — check if table style has borders
                    # If using "Table Grid" style, borders are visible by default
                    style_name = table.style.name if table.style else ''
                    if style_name == 'Normal Table':
                        # Normal Table has no borders by default — check cell borders
                        # Check if there's a vertical divider at cell level
                        cell_tc = table.cell(0, 0)._tc
                        tcPr = cell_tc.find(qn('w:tcPr'))
                        if tcPr is not None:
                            tcBorders = tcPr.find(qn('w:tcBorders'))
                            if tcBorders is not None:
                                right_b = tcBorders.find(qn('w:right'))
                                if right_b is not None and right_b.get(qn('w:val')) == 'single':
                                    sub_score += 0.25
                                    print("PASS: Component 4 — Normal Table with cell-level vertical divider")
                        if sub_score == 0:
                            print("FAIL: Component 4 — No border configuration found")
                    else:
                        print(f"FAIL: Component 4 — No tblBorders element, style={style_name}")
            else:
                print("FAIL: Component 4 — No tblPr element")

            if sub_score > 0:
                total_score += sub_score
            print(f"Component 4 total: {sub_score}/0.25 pts")
        else:
            print("FAIL: Component 4 — No table to check borders")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
