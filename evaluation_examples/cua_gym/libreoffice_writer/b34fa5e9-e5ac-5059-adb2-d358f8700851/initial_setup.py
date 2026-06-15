"""
Initial Setup: Supply chain report with Supplier Performance section (no table yet),
               and supplier_metrics.xlsx with Supplier 7 data.
Task ID: osworld_multi_apps_calc_to_writer_014
Domain: libreoffice_writer + libreoffice_calc (multi-app)
"""

import os
import shlex
import subprocess
import time

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
DOCS_DIR = '/home/user/Documents'
TASK_ID = 'osworld_multi_apps_calc_to_writer_014'
WRITER_OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
CALC_OUTPUT = f'{DOCS_DIR}/supplier_metrics.xlsx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_supplier_metrics():
    """Create ~/Documents/supplier_metrics.xlsx with realistic supplier data."""
    os.makedirs(DOCS_DIR, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Supplier Metrics'

    # Header row
    headers = ['Supplier ID', 'On-Time Rate', 'Order Accuracy', 'Returns', 'Rating']
    header_font = Font(name='Calibri', size=11, bold=True)
    header_fill = PatternFill(start_color='FF4472C4', end_color='FF4472C4', fill_type='solid')
    header_font_color = Font(name='Calibri', size=11, bold=True, color='FFFFFFFF')
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font_color
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    # Supplier data rows (realistic)
    data = [
        ['Supplier 1',  '94.2%', '98.5%', '1.2%', 4.7],
        ['Supplier 2',  '87.6%', '95.1%', '2.8%', 4.2],
        ['Supplier 3',  '91.3%', '97.4%', '1.5%', 4.5],
        ['Supplier 4',  '78.9%', '92.0%', '4.1%', 3.8],
        ['Supplier 5',  '96.8%', '99.2%', '0.7%', 4.9],
        ['Supplier 6',  '83.4%', '94.7%', '3.2%', 4.0],
        ['Supplier 7',  '88.5%', '96.3%', '2.1%', 4.3],
        ['Supplier 8',  '92.7%', '98.0%', '1.4%', 4.6],
        ['Supplier 9',  '75.1%', '89.5%', '5.3%', 3.5],
        ['Supplier 10', '90.0%', '96.8%', '1.9%', 4.4],
        ['Supplier 11', '85.3%', '93.6%', '3.0%', 4.1],
        ['Supplier 12', '97.2%', '99.5%', '0.5%', 5.0],
    ]

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
            if r % 2 == 0:
                cell.fill = PatternFill(start_color='FFD9E1F2', end_color='FFD9E1F2', fill_type='solid')

    # Column widths
    ws.column_dimensions['A'].width = 16
    ws.column_dimensions['B'].width = 14
    ws.column_dimensions['C'].width = 16
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 10
    ws.row_dimensions[1].height = 20

    wb.save(CALC_OUTPUT)
    print(f'Supplier metrics created: {CALC_OUTPUT}')


def create_writer_doc():
    """Create Writer document with supply chain report (Supplier Performance section, no table)."""
    doc = Document()

    # Title
    title = doc.add_heading('Supply Chain Performance Report', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / metadata
    subtitle = doc.add_paragraph('Q1 2025 | Procurement & Logistics Division')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report provides a comprehensive overview of our supply chain performance '
        'for Q1 2025. Key metrics including on-time delivery rates, order accuracy, '
        'return rates, and supplier ratings have been analyzed across all active '
        'procurement channels. Overall performance has remained stable with minor '
        'improvements in logistics efficiency compared to the previous quarter.'
    )

    doc.add_paragraph()

    # Section 2: Logistics Overview
    doc.add_heading('2. Logistics Overview', level=1)
    doc.add_paragraph(
        'Shipment volumes increased by 8.3% in Q1 2025 relative to Q4 2024. '
        'The average lead time across all suppliers was 4.2 days, down from '
        '4.8 days in the prior quarter. Regional distribution hubs in Southeast '
        'Asia demonstrated the highest throughput, processing over 14,000 units '
        'per week. Inventory turnover improved to 6.1x annually, exceeding the '
        'target of 5.8x set at the beginning of the fiscal year.'
    )
    doc.add_paragraph(
        'Freight costs were contained through renegotiated carrier contracts, '
        'resulting in a 5% reduction in per-unit shipping cost. Last-mile delivery '
        'performance improved significantly in urban markets, with 91% of orders '
        'delivered within the promised delivery window.'
    )

    doc.add_paragraph()

    # Section 3: Supplier Performance (key section — NO table yet)
    doc.add_heading('3. Supplier Performance', level=1)
    doc.add_paragraph(
        'The following section summarizes the performance metrics for all active '
        'suppliers during Q1 2025. Metrics are evaluated across four dimensions: '
        'on-time delivery rate, order accuracy, return rate, and overall supplier '
        'rating. Suppliers are ranked on a 5-point scale based on weighted '
        'performance criteria agreed upon in the annual supplier evaluation framework.'
    )

    doc.add_paragraph(
        'Data sourced from supplier_metrics.xlsx. Individual supplier records '
        'are available upon request from the procurement team.'
    )

    doc.add_paragraph()

    # Section 4: Risk Assessment
    doc.add_heading('4. Risk Assessment', level=1)
    doc.add_paragraph(
        'Three suppliers have been flagged for performance review due to sustained '
        'below-threshold metrics. Contingency sourcing plans have been activated '
        'for product categories with single-source dependency. The risk committee '
        'will convene in April to assess mitigation strategies for high-exposure '
        'commodity categories.'
    )

    doc.add_paragraph()

    # Section 5: Recommendations
    doc.add_heading('5. Recommendations', level=1)
    bullet_items = [
        'Renegotiate SLA terms with suppliers rated below 4.0 to include performance penalties.',
        'Expand dual-sourcing arrangements for critical components with lead times over 10 days.',
        'Implement automated performance dashboards to track real-time KPI deviations.',
        'Conduct quarterly business reviews with all Tier 1 suppliers.',
        'Evaluate green logistics options to reduce carbon footprint by 15% by end of 2025.',
    ]
    for item in bullet_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(WRITER_OUTPUT)
    print(f'Writer document created: {WRITER_OUTPUT}')


def main():
    create_supplier_metrics()
    create_writer_doc()

    # GUI-ready startup: open both files
    # First open the spreadsheet so it's accessible
    launch_gui(f'libreoffice --calc "{CALC_OUTPUT}"', delay_sec=2.0)
    # Then open the Writer document (this is the primary document the agent works in)
    launch_gui(f'libreoffice --writer "{WRITER_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Calc and Writer with DISPLAY=:0')


main()
