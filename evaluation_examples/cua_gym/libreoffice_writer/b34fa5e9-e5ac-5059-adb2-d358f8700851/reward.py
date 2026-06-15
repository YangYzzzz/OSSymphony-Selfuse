"""
Reward Script: Insert Supplier 7 delivery performance data as a table in Writer document
Task ID: osworld_multi_apps_calc_to_writer_014
Domain: libreoffice_writer (multi-apps: calc -> writer)
Scoring:
  Component 1: A table exists in the document (was not present in initial state)           — 0.3 pts
  Component 2: The table is located within the 'Supplier Performance' section              — 0.3 pts
  Component 3: The table contains correct Supplier 7 data (header + data row values)      — 0.4 pts
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_014'

# Expected Supplier 7 row values from supplier_metrics.xlsx
EXPECTED_SUPPLIER_7 = {
    'Supplier ID': 'Supplier 7',
    'On-Time Rate': '88.5%',
    'Order Accuracy': '96.3%',
    'Returns': '2.1%',
    'Rating': '4.3',
}

# Expected header columns (case-insensitive matching applied below)
EXPECTED_HEADERS = ['Supplier ID', 'On-Time Rate', 'Order Accuracy', 'Returns', 'Rating']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: A table was inserted into the document (0.3 points)
    # In the initial state the document has 0 tables.
    # A table must be present for any points to be awarded.
    # -----------------------------------------------------------------------
    try:
        num_tables = len(doc.tables)
        if num_tables >= 1:
            print(f"PASS: Component 1 — {num_tables} table(s) found in document (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 1 — No tables found in document (expected >= 1)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: The table is located within the 'Supplier Performance'
    # section (after the heading "3. Supplier Performance" and before
    # the heading "4. Risk Assessment") (0.3 points)
    # -----------------------------------------------------------------------
    try:
        body = doc.element.body
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # Collect body-level element descriptors in order: ('PARA', text) or ('TABLE',)
        body_elements = []
        for child in body:
            tag = child.tag
            if tag.endswith('}p'):
                text_parts = [t.text or '' for t in child.iter(f'{{{ns}}}t')]
                para_text = ''.join(text_parts).strip()
                body_elements.append(('PARA', para_text))
            elif tag.endswith('}tbl'):
                body_elements.append(('TABLE', ''))

        # Find the index of "3. Supplier Performance" heading and the next major section
        supplier_section_start = None
        next_section_start = None
        for idx, (etype, text) in enumerate(body_elements):
            if etype == 'PARA' and '3.' in text and 'Supplier Performance' in text:
                supplier_section_start = idx
            elif (supplier_section_start is not None
                  and next_section_start is None
                  and etype == 'PARA'
                  and (text.startswith('4.') or text.startswith('5.'))):
                next_section_start = idx

        # Count tables that fall between supplier section start and next section start
        tables_in_section = 0
        if supplier_section_start is not None:
            end = next_section_start if next_section_start is not None else len(body_elements)
            for etype, _ in body_elements[supplier_section_start:end]:
                if etype == 'TABLE':
                    tables_in_section += 1

        if tables_in_section >= 1:
            print("PASS: Component 2 — Table is located within the 'Supplier Performance' section (0.3 pts)")
            total_score += 0.3
        else:
            print("FAIL: Component 2 — No table found in the 'Supplier Performance' section")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: The table contains the correct Supplier 7 data row
    # (header row with column names + Supplier 7 data row values) (0.4 points)
    # -----------------------------------------------------------------------
    try:
        supplier7_table = None
        for table in doc.tables:
            if len(table.rows) >= 1:
                # Check if any row contains 'Supplier 7'
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells]
                    if any('Supplier 7' in t for t in row_texts):
                        supplier7_table = table
                        break
            if supplier7_table:
                break

        if supplier7_table is None:
            print("FAIL: Component 3 — No table with 'Supplier 7' data found")
        else:
            # Extract header row and data row
            rows = supplier7_table.rows
            header_row = [cell.text.strip() for cell in rows[0].cells]

            # Find the Supplier 7 data row
            data_row = None
            for row in rows:
                cells = [cell.text.strip() for cell in row.cells]
                if cells and 'Supplier 7' in cells[0]:
                    data_row = cells
                    break

            if data_row is None:
                print("FAIL: Component 3 — 'Supplier 7' not found in first column of any row")
            else:
                # Map header to data
                if len(header_row) != len(data_row):
                    print(f"FAIL: Component 3 — Column count mismatch: header={header_row}, data={data_row}")
                else:
                    # Build a dict from header -> value
                    row_dict = {h: v for h, v in zip(header_row, data_row)}

                    # Check Supplier ID
                    correct_id = row_dict.get('Supplier ID', '') == EXPECTED_SUPPLIER_7['Supplier ID']
                    # Check On-Time Rate
                    correct_on_time = row_dict.get('On-Time Rate', '') == EXPECTED_SUPPLIER_7['On-Time Rate']
                    # Check Order Accuracy
                    correct_accuracy = row_dict.get('Order Accuracy', '') == EXPECTED_SUPPLIER_7['Order Accuracy']
                    # Check Returns
                    correct_returns = row_dict.get('Returns', '') == EXPECTED_SUPPLIER_7['Returns']
                    # Check Rating (allow both '4.3' and 4.3 representation)
                    rating_val = row_dict.get('Rating', '')
                    correct_rating = str(rating_val).strip() == EXPECTED_SUPPLIER_7['Rating']

                    checks = {
                        'Supplier ID': correct_id,
                        'On-Time Rate': correct_on_time,
                        'Order Accuracy': correct_accuracy,
                        'Returns': correct_returns,
                        'Rating': correct_rating,
                    }

                    all_correct = all(checks.values())
                    failed = [k for k, v in checks.items() if not v]

                    print(f"  Data found: {row_dict}")
                    print(f"  Expected:   {EXPECTED_SUPPLIER_7}")
                    print(f"  Column checks: {checks}")

                    if all_correct:
                        print("PASS: Component 3 — Supplier 7 data row has all correct values (0.4 pts)")
                        total_score += 0.4
                    else:
                        print(f"FAIL: Component 3 — Incorrect values for: {failed}")

    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM environment
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
