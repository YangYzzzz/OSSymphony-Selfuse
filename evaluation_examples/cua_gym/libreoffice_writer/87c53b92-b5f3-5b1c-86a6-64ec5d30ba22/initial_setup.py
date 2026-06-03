"""
Initial Setup: Add Electronics inventory counts from Calc spreadsheet to Writer report
Task ID: osworld_multi_apps_calc_to_writer_009
Domain: libreoffice_writer (multi-app: calc + writer)
"""

import os
import shlex
import subprocess
import time
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_009'
DOCS_DIR = f'{WORKDIR}/Documents'
XLSX_OUTPUT = f'{DOCS_DIR}/inventory.xlsx'
DOCX_OUTPUT = f'{WORKDIR}/inventory_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_inventory_xlsx():
    """Create ~/Documents/inventory.xlsx with realistic multi-category inventory data."""
    os.makedirs(DOCS_DIR, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Inventory"

    # Header row
    headers = ['Category', 'Item Name', 'SKU', 'Quantity', 'Location']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, name='Calibri', size=11)
        cell.fill = PatternFill(start_color="FF4472C4", end_color="FF4472C4", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')

    # Inventory data — realistic multi-category content
    # Electronics items (8 items)
    data = [
        # Category,          Item Name,                SKU,           Quantity, Location
        ['Electronics',  '4K Smart TV 55"',          'ELEC-TV-001',      12,   'Warehouse A, Shelf 3'],
        ['Electronics',  'Wireless Noise-Canceling Headphones', 'ELEC-HC-002', 45, 'Warehouse A, Shelf 7'],
        ['Electronics',  'Laptop 15" Core i7',       'ELEC-LP-003',      8,    'Warehouse B, Shelf 1'],
        ['Electronics',  'Bluetooth Speaker Portable','ELEC-SP-004',     67,   'Warehouse A, Shelf 7'],
        ['Electronics',  'USB-C Charging Hub',        'ELEC-HB-005',    130,   'Warehouse B, Shelf 4'],
        ['Electronics',  'Digital Camera 24MP',       'ELEC-CM-006',     15,   'Warehouse A, Shelf 2'],
        ['Electronics',  'Tablet 10" 128GB',          'ELEC-TB-007',     23,   'Warehouse B, Shelf 1'],
        ['Electronics',  'Smart Watch Series 5',      'ELEC-WC-008',     38,   'Warehouse A, Shelf 5'],
        # Furniture items
        ['Furniture',    'Ergonomic Office Chair',    'FURN-CH-001',     20,   'Warehouse C, Bay 1'],
        ['Furniture',    'Standing Desk 60"',         'FURN-DK-002',      9,   'Warehouse C, Bay 2'],
        ['Furniture',    'Bookshelf 5-Tier',          'FURN-BS-003',     15,   'Warehouse C, Bay 3'],
        ['Furniture',    'Filing Cabinet 3-Drawer',   'FURN-FC-004',     11,   'Warehouse C, Bay 1'],
        # Office Supplies
        ['Office Supplies', 'A4 Copy Paper (500-sheet Ream)', 'OFFC-PP-001', 320, 'Warehouse D, Shelf 1'],
        ['Office Supplies', 'Black Ballpoint Pen Box/12',     'OFFC-PN-002', 180, 'Warehouse D, Shelf 2'],
        ['Office Supplies', 'Stapler Heavy Duty',             'OFFC-ST-003',  45, 'Warehouse D, Shelf 3'],
        ['Office Supplies', 'Whiteboard Markers Set/8',       'OFFC-MK-004',  72, 'Warehouse D, Shelf 2'],
        # Clothing
        ['Clothing',     'Corporate Polo Shirt (M)',  'CLTH-PS-001',     55,   'Warehouse E, Rack 1'],
        ['Clothing',     'High-Visibility Safety Vest','CLTH-VT-002',     40,   'Warehouse E, Rack 2'],
        # Sports
        ['Sports',       'Yoga Mat Non-Slip',         'SPRT-YM-001',     28,   'Warehouse F, Shelf 1'],
        ['Sports',       'Resistance Bands Set',      'SPRT-RB-002',     60,   'Warehouse F, Shelf 2'],
    ]

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = Font(name='Calibri', size=10)
            if c == 4:  # Quantity column — right-aligned
                cell.alignment = Alignment(horizontal='right')
            else:
                cell.alignment = Alignment(horizontal='left')

    # Column widths
    col_widths = [18, 38, 16, 12, 28]
    for col_idx, width in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = width

    # Auto-filter
    ws.auto_filter.ref = f"A1:E{len(data) + 1}"
    # Freeze header row
    ws.freeze_panes = "A2"

    wb.save(XLSX_OUTPUT)
    print(f'Inventory spreadsheet created: {XLSX_OUTPUT}')


def create_inventory_report_docx():
    """
    Create inventory_report.docx — a Writer document with a 'Current Stock Levels'
    section that does NOT yet contain any table (the agent must add the Electronics table).
    """
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('Quarterly Inventory Report', level=0)

    # Introduction paragraph
    intro = doc.add_paragraph(
        'This report provides a comprehensive overview of current inventory levels across '
        'all product categories. It is intended for use by the warehouse management team '
        'and procurement officers to monitor stock health and plan restocking activities.'
    )

    doc.add_paragraph('')  # spacing

    # Section 1: Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'As of the end of the current quarter, our total inventory spans five major product '
        'categories: Electronics, Furniture, Office Supplies, Clothing, and Sports Equipment. '
        'Overall stock levels remain within acceptable operational thresholds, though several '
        'high-demand electronics items are approaching reorder points.'
    )

    doc.add_paragraph('')

    # Section 2: Current Stock Levels — the agent inserts Electronics table here
    doc.add_heading('Current Stock Levels', level=1)
    doc.add_paragraph(
        'The table below should reflect the current inventory quantities for each category. '
        'Please refer to the inventory spreadsheet (~/Documents/inventory.xlsx) for the '
        'complete data set. Update the Electronics section with the latest figures from '
        'the spreadsheet.'
    )

    # Placeholder paragraph (no table — agent must add it)
    placeholder = doc.add_paragraph(
        '[Electronics inventory data to be inserted from inventory.xlsx]'
    )
    for run in placeholder.runs:
        run.italic = True
        run.font.color.rgb = __import__('docx').shared.RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph('')

    # Section 3: Reorder Alerts
    doc.add_heading('Reorder Alerts', level=1)
    doc.add_paragraph(
        'The following items have fallen below the minimum stock threshold and require '
        'immediate attention from the procurement team:'
    )
    alerts = [
        'Laptop 15" Core i7 (SKU: ELEC-LP-003) — Current: 8 units, Minimum: 10 units',
        'Standing Desk 60" (SKU: FURN-DK-002) — Current: 9 units, Minimum: 10 units',
        '4K Smart TV 55" (SKU: ELEC-TV-001) — Current: 12 units, Minimum: 15 units',
    ]
    for alert in alerts:
        doc.add_paragraph(alert, style='List Bullet')

    doc.add_paragraph('')

    # Section 4: Notes
    doc.add_heading('Notes & Observations', level=1)
    doc.add_paragraph(
        'All inventory counts were verified during the last physical stock-take conducted '
        'on the first Monday of this month. Discrepancies between system records and '
        'physical counts have been escalated to the inventory control supervisor. '
        'Next scheduled stock-take: end of quarter.'
    )

    doc.save(DOCX_OUTPUT)
    print(f'Inventory report document created: {DOCX_OUTPUT}')


def create_initial():
    create_inventory_xlsx()
    create_inventory_report_docx()

    # GUI-ready startup:
    # 1. Open the inventory report in LibreOffice Writer (primary task window)
    launch_gui(f'libreoffice --writer "{DOCX_OUTPUT}"', delay_sec=2.0)
    # 2. Open the inventory spreadsheet in LibreOffice Calc (data source)
    launch_gui(f'libreoffice --calc "{XLSX_OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer (inventory_report.docx) and Calc (inventory.xlsx) with DISPLAY=:0')


create_initial()
