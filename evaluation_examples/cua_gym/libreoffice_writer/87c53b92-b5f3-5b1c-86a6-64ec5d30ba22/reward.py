"""
Reward Script: Add Electronics inventory counts to inventory_report.docx
Task ID: osworld_multi_apps_calc_to_writer_009
Domain: libreoffice_writer (multi-app: calc + writer)

Scoring Rubric:
  Component 1 (0.3 pts): A table exists in the document body within the 'Current Stock Levels' section
  Component 2 (0.4 pts): Table contains all 8 Electronics items with correct quantities
  Component 3 (0.3 pts): Table has correct structure (headers: Item Name, SKU, Quantity, Location)
  Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_009'

# Ground truth Electronics items from inventory.xlsx
EXPECTED_ELECTRONICS = [
    ('4K Smart TV 55"',                   'ELEC-TV-001',  12),
    ('Wireless Noise-Canceling Headphones','ELEC-HC-002',  45),
    ('Laptop 15" Core i7',                'ELEC-LP-003',   8),
    ('Bluetooth Speaker Portable',         'ELEC-SP-004',  67),
    ('USB-C Charging Hub',                 'ELEC-HB-005', 130),
    ('Digital Camera 24MP',               'ELEC-CM-006',  15),
    ('Tablet 10" 128GB',                  'ELEC-TB-007',  23),
    ('Smart Watch Series 5',              'ELEC-WC-008',  38),
]

# Expected column headers (task requires Item Name and Quantity at minimum;
# full table has Item Name, SKU, Quantity, Location)
EXPECTED_HEADERS = {'item name', 'sku', 'quantity', 'location'}
REQUIRED_HEADERS = {'item name', 'quantity'}


def get_table_after_current_stock_levels(doc):
    """
    Find the first table that appears after the 'Current Stock Levels' heading
    in document body order. Returns the table element if found, else None.
    """
    body = doc.element.body
    found_heading = False
    for elem in body:
        tag = elem.tag.split('}')[1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            text = ''.join(r.text for r in elem.iter(qn('w:t')) if r.text)
            if 'current stock levels' in text.lower():
                found_heading = True
        elif tag == 'tbl' and found_heading:
            return elem
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: verify the document has the expected headings
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    if 'Current Stock Levels' not in headings:
        print("FAIL: 'Current Stock Levels' heading not found — document structure is invalid")
        print("REWARD: 0.0")
        return 0.0

    # ------------------------------------------------------------------
    # Component 1: A table exists in the document and is placed within
    #              the 'Current Stock Levels' section (0.3 points)
    # ------------------------------------------------------------------
    try:
        tbl_elem = get_table_after_current_stock_levels(doc)
        tables_in_doc = len(doc.tables)

        if tbl_elem is not None and tables_in_doc >= 1:
            print(f"PASS: Component 1 — Table found in 'Current Stock Levels' section "
                  f"(total tables in doc: {tables_in_doc}) (0.3 pts)")
            total_score += 0.3
        elif tables_in_doc >= 1:
            print(f"FAIL: Component 1 — Table exists in document but NOT after "
                  f"'Current Stock Levels' heading")
        else:
            print(f"FAIL: Component 1 — No table found in document (found {tables_in_doc} tables)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ------------------------------------------------------------------
    # Component 2: Table contains all 8 Electronics items with correct
    #              quantities (0.4 points)
    # ------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 — No table in document to inspect")
        else:
            # Use the first table in the document
            table = doc.tables[0]
            all_cells_text = []
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells]
                all_cells_text.append(row_texts)

            # Build flat text of all cells for item/quantity lookup
            flat_cell_text = ' '.join(
                cell.text.strip().lower()
                for row in table.rows
                for cell in row.cells
            )

            items_found = 0
            qty_correct = 0

            for item_name, sku, qty in EXPECTED_ELECTRONICS:
                item_name_lower = item_name.lower()
                # Check if item name appears in any cell
                item_present = any(
                    item_name_lower in cell.text.strip().lower()
                    for row in table.rows
                    for cell in row.cells
                )
                if item_present:
                    items_found += 1
                    # Check quantity appears in same row as item
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        row_combined = ' '.join(row_text).lower()
                        if item_name_lower in row_combined:
                            # Quantity should appear somewhere in the row
                            qty_strs = [str(qty), str(float(qty))]
                            if any(q in row_text for q in qty_strs):
                                qty_correct += 1
                            break

            total_items = len(EXPECTED_ELECTRONICS)
            print(f"INFO: Component 2 — Items found: {items_found}/{total_items}, "
                  f"quantities correct: {qty_correct}/{total_items}")

            if items_found == total_items and qty_correct == total_items:
                print(f"PASS: Component 2 — All {total_items} Electronics items with correct quantities (0.4 pts)")
                total_score += 0.4
            elif items_found == total_items and qty_correct >= total_items * 0.75:
                # Partial: all items present but some quantities wrong or format issue
                print(f"PASS (partial): Component 2 — All items present, "
                      f"{qty_correct}/{total_items} quantities correct (0.2 pts)")
                total_score += 0.2
            elif items_found >= total_items * 0.75:
                print(f"FAIL (partial): Component 2 — Only {items_found}/{total_items} items found")
            else:
                print(f"FAIL: Component 2 — Only {items_found}/{total_items} Electronics items found in table")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ------------------------------------------------------------------
    # Component 3: Table has correct column headers (Item Name, SKU,
    #              Quantity, Location) or at minimum Item Name + Quantity
    #              (0.3 points)
    # ------------------------------------------------------------------
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 — No table in document")
        else:
            table = doc.tables[0]
            # First row is expected to be headers
            first_row_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
            first_row_set = set(first_row_texts)

            has_required = REQUIRED_HEADERS.issubset(first_row_set)
            has_full_headers = EXPECTED_HEADERS.issubset(first_row_set)

            if has_full_headers:
                print(f"PASS: Component 3 — Full headers found: {first_row_texts} (0.3 pts)")
                total_score += 0.3
            elif has_required:
                print(f"PASS (partial): Component 3 — Required headers 'item name' and 'quantity' "
                      f"found (partial: {first_row_texts}) (0.15 pts)")
                total_score += 0.15
            else:
                # Check if table has any data row that looks like a header (case-insensitive partial match)
                all_row_texts = []
                for row in table.rows:
                    all_row_texts.append([cell.text.strip().lower() for cell in row.cells])

                header_like = any(
                    any('item' in cell or 'quantity' in cell or 'sku' in cell
                        for cell in row)
                    for row in all_row_texts
                )

                if header_like:
                    print(f"FAIL (partial): Component 3 — Header-like row found but missing "
                          f"required columns. First row: {first_row_texts}")
                else:
                    print(f"FAIL: Component 3 — No recognizable headers in table. "
                          f"First row: {first_row_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.1f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


# Entry point: test against the canonical artifact path on the VM
file_path = f'{WORKDIR}/inventory_report.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
