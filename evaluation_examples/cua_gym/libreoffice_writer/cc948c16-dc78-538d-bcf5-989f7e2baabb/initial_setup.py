"""
Initial Setup: Create a Writer document with command-line examples in default style.
Task ID: writer_tech_043
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_043'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    heading = doc.add_heading('Linux Command Reference Guide', level=1)

    # --- Introductory paragraph ---
    doc.add_paragraph(
        'This document provides a quick reference for commonly used Linux '
        'terminal commands. Each section includes practical examples that '
        'system administrators and developers encounter in daily workflows.'
    )

    # --- Section 1: File Operations ---
    doc.add_heading('File and Directory Operations', level=2)
    doc.add_paragraph(
        'The following commands are essential for navigating the file system '
        'and managing files on a Linux workstation.'
    )

    # Terminal examples in DEFAULT style (no special formatting)
    doc.add_paragraph('ls -la /var/log/')
    doc.add_paragraph('find /home -name "*.conf" -type f -mtime -7')
    doc.add_paragraph('cp -rv /etc/nginx/sites-available/ /backup/nginx/')
    doc.add_paragraph('chmod 755 /opt/scripts/deploy.sh')

    doc.add_paragraph(
        'These commands help locate configuration files, create backups, '
        'and set appropriate permissions on deployment scripts.'
    )

    # --- Section 2: Process Management ---
    doc.add_heading('Process Management', level=2)
    doc.add_paragraph(
        'Monitoring and controlling processes is a critical part of server '
        'administration. Below are commands frequently used to inspect and '
        'manage running services.'
    )

    doc.add_paragraph('ps aux | grep nginx')
    doc.add_paragraph('top -b -n 1 | head -20')
    doc.add_paragraph('kill -SIGTERM $(pidof apache2)')
    doc.add_paragraph('systemctl restart postgresql.service')

    doc.add_paragraph(
        'Always verify that services have restarted correctly by checking '
        'their status with systemctl status.'
    )

    # --- Section 3: Networking ---
    doc.add_heading('Network Diagnostics', level=2)
    doc.add_paragraph(
        'Network troubleshooting requires a combination of tools to identify '
        'connectivity issues and monitor traffic.'
    )

    doc.add_paragraph('curl -I https://api.example.com/health')
    doc.add_paragraph('netstat -tlnp | grep :8080')
    doc.add_paragraph('ssh -L 3306:db-server:3306 bastion@10.0.1.50')
    doc.add_paragraph('tcpdump -i eth0 -w /tmp/capture.pcap port 443')

    doc.add_paragraph(
        'When diagnosing production issues, always capture traffic on the '
        'correct interface and filter by port to reduce noise.'
    )

    # --- Section 4: Disk and Storage ---
    doc.add_heading('Disk and Storage', level=2)
    doc.add_paragraph(
        'Monitoring disk usage prevents outages caused by full filesystems. '
        'These commands help identify space consumption patterns.'
    )

    doc.add_paragraph('df -h /dev/sda1')
    doc.add_paragraph('du -sh /var/log/* | sort -rh | head -10')
    doc.add_paragraph('lsblk -f')

    doc.add_paragraph(
        'Schedule regular disk usage reports using cron jobs to proactively '
        'identify storage growth before it becomes critical.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
