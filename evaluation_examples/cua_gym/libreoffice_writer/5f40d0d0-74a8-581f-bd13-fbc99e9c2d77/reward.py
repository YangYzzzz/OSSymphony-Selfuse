"""
Reward Script: Insert 'DRAFT - Internal Use Only' as header on every page
Task ID: writer_tech_012
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Header is enabled (not empty) in all sections
  Component 2 (0.6): Header text matches 'DRAFT - Internal Use Only' in all sections
"""

import os
import time


WORKDIR = '/home/user'
TASK_ID = 'writer_tech_012'


def persist_app_state(domain: str):
    """Save any unsaved changes in LibreOffice before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)
    print(f"INFO: Document has {num_sections} section(s)")

    # Component 1: Header is enabled with non-empty text in all sections (0.4 points)
    # Initial env has empty header text; golden env has header text set.
    try:
        sections_with_header = 0
        for i, section in enumerate(doc.sections):
            header = section.header
            # Collect all text from all header paragraphs
            header_text = ''.join(p.text for p in header.paragraphs).strip()
            if header_text:
                sections_with_header += 1
                print(f"  Section {i}: header text present ('{header_text[:50]}')")
            else:
                print(f"  Section {i}: header text is EMPTY")

        if num_sections > 0 and sections_with_header == num_sections:
            print(f"PASS: Component 1 -- Header enabled in all {num_sections} section(s) (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 -- Header enabled in {sections_with_header}/{num_sections} sections")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Header text is exactly 'DRAFT - Internal Use Only' in all sections (0.6 points)
    # This is the core task requirement.
    EXPECTED_TEXT = 'DRAFT - Internal Use Only'
    try:
        sections_with_correct_text = 0
        for i, section in enumerate(doc.sections):
            header = section.header
            header_text = ''.join(p.text for p in header.paragraphs).strip()
            if header_text == EXPECTED_TEXT:
                sections_with_correct_text += 1
                print(f"  Section {i}: header text matches expected")
            else:
                print(f"  Section {i}: header text mismatch: expected '{EXPECTED_TEXT}', found '{header_text}'")

        if num_sections > 0 and sections_with_correct_text == num_sections:
            print(f"PASS: Component 2 -- Correct header text in all {num_sections} section(s) (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 2 -- Correct text in {sections_with_correct_text}/{num_sections} sections")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
