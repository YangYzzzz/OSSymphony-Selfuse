"""
Initial Setup: Insert header on API specification document
Task ID: writer_tech_012
Domain: libreoffice_writer

Creates a 6-page API specification document with NO header.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_012'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add an explicit page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Ensure NO header is set (default is no header, but be explicit)
    header = section.header
    header.is_linked_to_previous = True

    # ========== PAGE 1: Title & Overview ==========
    title = doc.add_heading('Nexus Platform API Specification', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Version 3.2.1 | Last Updated: March 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'The Nexus Platform API provides programmatic access to core platform '
        'services including user management, data analytics, workflow automation, '
        'and notification delivery. This document specifies all available endpoints, '
        'request/response schemas, authentication requirements, and rate limiting policies.'
    )
    doc.add_paragraph(
        'All API endpoints are served over HTTPS at the base URL '
        'https://api.nexusplatform.io/v3. Requests must include a valid Bearer token '
        'in the Authorization header. Responses are returned in JSON format with '
        'appropriate HTTP status codes.'
    )

    doc.add_heading('1.1 Authentication', level=2)
    doc.add_paragraph(
        'Authentication is handled via OAuth 2.0 with JWT tokens. Clients must first '
        'obtain an access token by presenting valid credentials to the /auth/token endpoint. '
        'Tokens expire after 3600 seconds and can be refreshed using the refresh_token grant type.'
    )

    # Auth table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['Parameter', 'Type', 'Description']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    auth_data = [
        ['grant_type', 'string', 'Must be "client_credentials" or "refresh_token"'],
        ['client_id', 'string', 'Application client identifier (UUID format)'],
        ['client_secret', 'string', 'Application secret key (min 32 characters)'],
    ]
    for r, row_data in enumerate(auth_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    add_page_break(doc)

    # ========== PAGE 2: User Management Endpoints ==========
    doc.add_heading('2. User Management Endpoints', level=1)

    doc.add_heading('2.1 GET /users', level=2)
    doc.add_paragraph(
        'Retrieves a paginated list of users in the organization. Supports filtering '
        'by role, department, and account status. Results are sorted by creation date '
        'in descending order by default.'
    )

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    h2 = ['Parameter', 'Type', 'Required', 'Description']
    for i, h in enumerate(h2):
        run = table2.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    users_params = [
        ['page', 'integer', 'No', 'Page number (default: 1)'],
        ['per_page', 'integer', 'No', 'Items per page (default: 25, max: 100)'],
        ['role', 'string', 'No', 'Filter by role: admin, editor, viewer'],
        ['department', 'string', 'No', 'Filter by department code (e.g., ENG, MKT)'],
        ['status', 'string', 'No', 'Filter by status: active, suspended, pending'],
    ]
    for r, row_data in enumerate(users_params, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_heading('2.2 POST /users', level=2)
    doc.add_paragraph(
        'Creates a new user account. The requesting client must have admin privileges. '
        'An invitation email is sent automatically upon successful creation. The response '
        'includes the new user ID and a temporary access link valid for 72 hours.'
    )

    doc.add_heading('2.3 PUT /users/{user_id}', level=2)
    doc.add_paragraph(
        'Updates an existing user profile. Partial updates are supported; only the fields '
        'included in the request body will be modified. The user_id path parameter must be '
        'a valid UUID. Attempting to modify a suspended account returns HTTP 403.'
    )

    add_page_break(doc)

    # ========== PAGE 3: Data Analytics Endpoints ==========
    doc.add_heading('3. Data Analytics Endpoints', level=1)

    doc.add_heading('3.1 POST /analytics/query', level=2)
    doc.add_paragraph(
        'Executes an analytics query against the platform data warehouse. Queries use '
        'NexusQL syntax, a SQL-like language supporting aggregations, window functions, '
        'and time-series operations. Maximum query execution time is 120 seconds.'
    )

    doc.add_paragraph(
        'The request body must include a "query" field containing the NexusQL statement '
        'and an optional "parameters" object for parameterized queries. Results are returned '
        'as an array of row objects with column names as keys.'
    )

    table3 = doc.add_table(rows=5, cols=4)
    table3.style = 'Table Grid'
    h3 = ['Field', 'Type', 'Required', 'Description']
    for i, h in enumerate(h3):
        run = table3.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    analytics_data = [
        ['query', 'string', 'Yes', 'NexusQL query statement'],
        ['parameters', 'object', 'No', 'Named parameter bindings for the query'],
        ['timeout', 'integer', 'No', 'Custom timeout in seconds (max: 120)'],
        ['format', 'string', 'No', 'Response format: json (default), csv, parquet'],
    ]
    for r, row_data in enumerate(analytics_data, 1):
        for c, val in enumerate(row_data):
            table3.cell(r, c).text = val

    doc.add_heading('3.2 GET /analytics/reports/{report_id}', level=2)
    doc.add_paragraph(
        'Retrieves a previously generated analytics report by its unique identifier. '
        'Reports are cached for 24 hours after initial generation. Expired reports '
        'return HTTP 410 Gone with a message indicating the report must be regenerated.'
    )

    doc.add_heading('3.3 GET /analytics/dashboards', level=2)
    doc.add_paragraph(
        'Lists all dashboards accessible to the authenticated user. Each dashboard '
        'object includes its title, widget configuration, refresh interval, and sharing '
        'permissions. Dashboards are returned sorted by last-modified date.'
    )

    add_page_break(doc)

    # ========== PAGE 4: Workflow Automation ==========
    doc.add_heading('4. Workflow Automation Endpoints', level=1)

    doc.add_heading('4.1 POST /workflows', level=2)
    doc.add_paragraph(
        'Creates a new workflow definition. Workflows consist of a series of steps, '
        'each with a trigger condition and one or more actions. Steps execute sequentially '
        'unless parallel execution is explicitly enabled via the "parallel" flag.'
    )
    doc.add_paragraph(
        'Supported trigger types include: schedule (cron expression), webhook (HTTP callback), '
        'event (platform event bus), and manual (user-initiated). Each trigger type has its '
        'own configuration schema documented in Section 4.3.'
    )

    doc.add_heading('4.2 GET /workflows/{workflow_id}/executions', level=2)
    doc.add_paragraph(
        'Returns the execution history for a specific workflow. Each execution record '
        'includes the start time, end time, status (running, completed, failed, cancelled), '
        'step-level logs, and any error messages. Pagination is supported with default '
        'page size of 50 records.'
    )

    table4 = doc.add_table(rows=5, cols=3)
    table4.style = 'Table Grid'
    h4 = ['Status Code', 'Meaning', 'Common Cause']
    for i, h in enumerate(h4):
        run = table4.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    status_data = [
        ['200', 'Success', 'Request processed successfully'],
        ['400', 'Bad Request', 'Invalid workflow definition or missing required fields'],
        ['404', 'Not Found', 'Workflow ID does not exist or has been deleted'],
        ['429', 'Too Many Requests', 'Rate limit exceeded (100 requests per minute)'],
    ]
    for r, row_data in enumerate(status_data, 1):
        for c, val in enumerate(row_data):
            table4.cell(r, c).text = val

    doc.add_heading('4.3 Trigger Configuration Schemas', level=2)
    doc.add_paragraph(
        'Schedule triggers accept a standard cron expression with optional timezone. '
        'Webhook triggers require a callback URL and an optional HMAC secret for signature '
        'verification. Event triggers bind to one or more event types from the platform bus.'
    )

    add_page_break(doc)

    # ========== PAGE 5: Notification Service ==========
    doc.add_heading('5. Notification Service Endpoints', level=1)

    doc.add_heading('5.1 POST /notifications/send', level=2)
    doc.add_paragraph(
        'Sends a notification to one or more recipients through the specified channel. '
        'Supported channels include email, SMS, push notification, and in-app messaging. '
        'Bulk sends are limited to 1000 recipients per request.'
    )
    doc.add_paragraph(
        'Each notification requires a template_id referencing a pre-configured template '
        'and a data object containing the merge fields. Templates support conditional '
        'sections, loops, and localization with automatic fallback to the default locale.'
    )

    doc.add_heading('5.2 GET /notifications/templates', level=2)
    doc.add_paragraph(
        'Lists all notification templates available in the organization. Templates are '
        'versioned; the latest published version is used by default unless a specific '
        'version number is provided in the send request.'
    )

    table5 = doc.add_table(rows=6, cols=3)
    table5.style = 'Table Grid'
    h5 = ['Channel', 'Max Length', 'Rate Limit']
    for i, h in enumerate(h5):
        run = table5.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    notif_data = [
        ['Email', 'No limit', '500 per hour per sender'],
        ['SMS', '160 characters', '100 per hour per recipient'],
        ['Push', '4096 bytes payload', '1000 per hour per app'],
        ['In-App', '10,000 characters', '5000 per hour per org'],
        ['Slack', '40,000 characters', '200 per hour per channel'],
    ]
    for r, row_data in enumerate(notif_data, 1):
        for c, val in enumerate(row_data):
            table5.cell(r, c).text = val

    doc.add_heading('5.3 Delivery Status Webhooks', level=2)
    doc.add_paragraph(
        'The notification service supports delivery status callbacks via webhooks. '
        'Register a webhook URL through the /notifications/webhooks endpoint to receive '
        'real-time delivery status updates including: delivered, bounced, opened, and clicked events.'
    )

    add_page_break(doc)

    # ========== PAGE 6: Rate Limiting & Error Handling ==========
    doc.add_heading('6. Rate Limiting & Error Handling', level=1)

    doc.add_heading('6.1 Rate Limits', level=2)
    doc.add_paragraph(
        'All API endpoints are subject to rate limiting. The default rate limit is '
        '1000 requests per minute per API key. Premium tier clients receive a limit '
        'of 5000 requests per minute. Rate limit headers are included in every response.'
    )

    table6 = doc.add_table(rows=4, cols=3)
    table6.style = 'Table Grid'
    h6 = ['Header', 'Type', 'Description']
    for i, h in enumerate(h6):
        run = table6.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True
    rate_data = [
        ['X-RateLimit-Limit', 'integer', 'Maximum requests allowed per window'],
        ['X-RateLimit-Remaining', 'integer', 'Requests remaining in current window'],
        ['X-RateLimit-Reset', 'timestamp', 'UTC epoch when the window resets'],
    ]
    for r, row_data in enumerate(rate_data, 1):
        for c, val in enumerate(row_data):
            table6.cell(r, c).text = val

    doc.add_heading('6.2 Error Response Format', level=2)
    doc.add_paragraph(
        'All error responses follow a consistent JSON structure with the following fields: '
        '"error" (string error code), "message" (human-readable description), "details" '
        '(optional array of field-level errors), and "request_id" (UUID for support reference).'
    )

    doc.add_heading('6.3 Retry Strategy', level=2)
    doc.add_paragraph(
        'Clients should implement exponential backoff for retryable errors (HTTP 429, 500, 502, '
        '503, 504). The recommended initial delay is 1 second with a multiplier of 2 and a '
        'maximum of 5 retries. The Retry-After header, when present, should be honored over '
        'the calculated backoff interval.'
    )

    doc.add_heading('6.4 Deprecation Policy', level=2)
    doc.add_paragraph(
        'Deprecated endpoints are announced at least 6 months before removal. The '
        'Sunset header indicates the deprecation date. Deprecated endpoints return '
        'a Warning header with code 299 and a deprecation notice message.'
    )

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
