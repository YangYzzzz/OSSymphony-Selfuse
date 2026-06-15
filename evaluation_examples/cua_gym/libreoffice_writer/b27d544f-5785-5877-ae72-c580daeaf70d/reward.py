"""
Reward Script: Line spacing variation by paragraph type in a newsletter
Task ID: osworld_writer_line_spacing_per_paragraph_010
Domain: libreoffice_writer
Scoring:
  Component 1 (0.6 pts): All 3 intro paragraphs (doc paras 2, 4, 6 = 0-based indices 1, 3, 5) have double spacing (2.0)
  Component 2 (0.4 pts): The spacing contrast is correct: intro paras are double AND continuation paras (3, 5, 7 = indices 2, 4, 6) are single spacing (1.0)
                         This also ensures continuation paras were not accidentally doubled.
                         Both conditions must hold for this component to pass; since intro paras must be double (fails on initial),
                         this component also correctly fails on the initial state.
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_010'


def persist_app_state():
    """Send ctrl+s to save any unsaved LibreOffice edits."""
    try:
        import pyautogui
        import time
        os.environ["DISPLAY"] = ":0"
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    The document has 7 paragraphs:
      - Para 1 (index 0): Title (should be unchanged)
      - Paras 2, 4, 6 (indices 1, 3, 5): Section intro paragraphs — MUST have double spacing (2.0)
      - Paras 3, 5, 7 (indices 2, 4, 6): Section continuation paragraphs — MUST remain single spacing (1.0)

    Scoring is designed so that initial_env returns 0.0 (intro paras are single in initial state).

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Pre-condition gate: document must have at least 7 paragraphs
    if len(doc.paragraphs) < 7:
        print(f"CRITICAL: Expected at least 7 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    paras = doc.paragraphs  # 0-indexed list
    intro_indices = [1, 3, 5]    # 0-based: doc paragraphs 2, 4, 6 (intro/first of each section)
    cont_indices  = [2, 4, 6]    # 0-based: doc paragraphs 3, 5, 7 (continuation/second of each section)

    # -------------------------------------------------------------------------
    # Component 1 (0.6 pts): Intro paragraphs (2, 4, 6) have double spacing
    # This FAILS on initial_env (all paras are single spacing there)
    # This PASSES on golden_env (intro paras are set to 2.0)
    # -------------------------------------------------------------------------
    try:
        intro_results = []
        for idx in intro_indices:
            p = paras[idx]
            ls = p.paragraph_format.line_spacing
            is_double = (ls is not None and abs(float(ls) - 2.0) < 0.05)
            intro_results.append(is_double)
            print(f"  Intro para {idx+1} (0-based {idx}): line_spacing={ls}, is_double={is_double}, "
                  f"text_preview={repr(p.text[:40])}")

        if all(intro_results):
            print("PASS: Component 1 — All 3 intro paragraphs (2,4,6) have double spacing (0.6 pts)")
            total_score += 0.6
        elif any(intro_results):
            count = sum(intro_results)
            partial = round(count / 3 * 0.6, 4)
            if partial > 0:
                print(f"PARTIAL: Component 1 — {count}/3 intro paragraphs have double spacing ({partial} pts)")
                total_score += partial
        else:
            print("FAIL: Component 1 — No intro paragraphs have double spacing. Expected 2.0 for doc paras 2, 4, 6.")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2 (0.4 pts): Correct spacing contrast — intro paras are double
    #   AND continuation paras remain single spacing.
    # This is a compound check. Because it requires intro paras to be double,
    # it also FAILS on initial_env (intro paras are single there).
    # -------------------------------------------------------------------------
    try:
        # Re-check intro paras (double)
        intro_double_ok = all(
            paras[idx].paragraph_format.line_spacing is not None and
            abs(float(paras[idx].paragraph_format.line_spacing) - 2.0) < 0.05
            for idx in intro_indices
        )
        # Check continuation paras (single: 1.0 or None means inherit single)
        from docx.enum.text import WD_LINE_SPACING
        cont_single_ok = all(
            (paras[idx].paragraph_format.line_spacing is None) or
            (paras[idx].paragraph_format.line_spacing is not None and
             abs(float(paras[idx].paragraph_format.line_spacing) - 1.0) < 0.05) or
            (paras[idx].paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE)
            for idx in cont_indices
        )

        for idx in cont_indices:
            p = paras[idx]
            ls = p.paragraph_format.line_spacing
            rule = p.paragraph_format.line_spacing_rule
            print(f"  Cont para {idx+1} (0-based {idx}): line_spacing={ls}, rule={rule}, "
                  f"text_preview={repr(p.text[:40])}")

        if intro_double_ok and cont_single_ok:
            print("PASS: Component 2 — Correct spacing contrast: intro=double, continuation=single (0.4 pts)")
            total_score += 0.4
        elif intro_double_ok and not cont_single_ok:
            print("FAIL: Component 2 — Intro paras are double but continuation paras were altered (not single).")
        elif not intro_double_ok and cont_single_ok:
            print("FAIL: Component 2 — Continuation paras are single but intro paras are not double.")
        else:
            print("FAIL: Component 2 — Neither intro paras are double nor continuation paras are correctly single.")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {round(total_score, 4)}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
