"""
Initial Setup: Line spacing varies by paragraph type in a newsletter
Task ID: osworld_writer_line_spacing_per_paragraph_010
Domain: libreoffice_writer

Creates a company newsletter document with 3 sections, each having 2 paragraphs.
All paragraphs start at single spacing. The agent must:
  - Set introduction paragraphs (1, 3, 5) to double spacing
  - Keep continuation paragraphs (2, 4, 6) at single spacing
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_010'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Document Title ---
    title = doc.add_paragraph("Brightfield Technologies — Q1 2025 Employee Newsletter")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)

    # --- Section 1 ---
    # Section 1 intro paragraph (paragraph index 0 in content paragraphs - will become double)
    p1_intro = doc.add_paragraph(
        "Welcome to the first quarter update from Brightfield Technologies. "
        "This has been a remarkable period of growth and innovation across all departments. "
        "We are proud to share the accomplishments and milestones that define our team's dedication. "
        "As we move forward into Q2, the momentum we have built will continue to drive our success."
    )
    p1_intro.paragraph_format.line_spacing = 1.0
    p1_intro.paragraph_format.space_before = Pt(6)
    p1_intro.paragraph_format.space_after = Pt(6)

    # Section 1 continuation paragraph (paragraph index 1 - stays single)
    p1_cont = doc.add_paragraph(
        "In January, the engineering division completed the deployment of the new cloud infrastructure "
        "that had been in planning since last October. The transition was seamless, reducing system "
        "downtime by 42% compared to the previous year. The DevOps team, led by Priya Ramaswamy, "
        "coordinated more than 120 individual migration tasks over a six-week period."
    )
    p1_cont.paragraph_format.line_spacing = 1.0
    p1_cont.paragraph_format.space_before = Pt(6)
    p1_cont.paragraph_format.space_after = Pt(12)

    # --- Section 2 ---
    # Section 2 intro paragraph (paragraph index 2 - will become double)
    p2_intro = doc.add_paragraph(
        "Our sales and customer success teams have achieved outstanding results in this quarter. "
        "The combined efforts of over 60 account managers and support specialists culminated in a "
        "record-breaking revenue month in March. Customer satisfaction scores reached an all-time "
        "high of 94%, reflecting the hard work of everyone involved in client engagement."
    )
    p2_intro.paragraph_format.line_spacing = 1.0
    p2_intro.paragraph_format.space_before = Pt(6)
    p2_intro.paragraph_format.space_after = Pt(6)

    # Section 2 continuation paragraph (paragraph index 3 - stays single)
    p2_cont = doc.add_paragraph(
        "Among the quarter's highlights, the new enterprise partnership with Solaris Financial Group "
        "stands out as a major milestone. Negotiated by Senior Account Executive Daniel Okonkwo, "
        "the contract spans three years and is valued at approximately $2.4 million. This agreement "
        "opens the door to a new vertical market that the company has been targeting for two years."
    )
    p2_cont.paragraph_format.line_spacing = 1.0
    p2_cont.paragraph_format.space_before = Pt(6)
    p2_cont.paragraph_format.space_after = Pt(12)

    # --- Section 3 ---
    # Section 3 intro paragraph (paragraph index 4 - will become double)
    p3_intro = doc.add_paragraph(
        "The people and culture team has introduced a series of new initiatives designed to support "
        "employee wellbeing and professional development throughout the coming year. These programs "
        "reflect our belief that a thriving workforce is the foundation of long-term business success. "
        "We encourage all staff members to take advantage of the resources now available to them."
    )
    p3_intro.paragraph_format.line_spacing = 1.0
    p3_intro.paragraph_format.space_before = Pt(6)
    p3_intro.paragraph_format.space_after = Pt(6)

    # Section 3 continuation paragraph (paragraph index 5 - stays single)
    p3_cont = doc.add_paragraph(
        "Beginning in April, the company will offer subsidised access to an online learning platform "
        "featuring more than 5,000 courses across technology, leadership, and creative disciplines. "
        "Additionally, a new mentorship programme pairs junior employees with senior colleagues for "
        "structured monthly check-ins over a six-month engagement cycle managed by HR."
    )
    p3_cont.paragraph_format.line_spacing = 1.0
    p3_cont.paragraph_format.space_before = Pt(6)
    p3_cont.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
