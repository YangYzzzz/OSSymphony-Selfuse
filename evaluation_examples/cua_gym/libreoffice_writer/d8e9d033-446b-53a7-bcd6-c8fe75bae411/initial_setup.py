"""
Initial Setup: Add AutoCorrect entry for 'addr'
Task ID: writer_edit_054
Domain: libreoffice_writer

Creates a business letter template at ~/Desktop/letter_template.docx
and opens it in LibreOffice Writer.
The AutoCorrect entry for 'addr' does NOT exist yet (that is the task).
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_054'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/letter_template.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Company letterhead
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run("Nexus Technology Solutions")
    run.bold = True
    run.font.size = Pt(16)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tagline.add_run("Empowering Innovation Worldwide")

    # Horizontal line placeholder (blank paragraph for spacing)
    doc.add_paragraph()

    # Date line
    date_para = doc.add_paragraph()
    date_para.add_run("Date: March 15, 2025")

    doc.add_paragraph()

    # Recipient block
    doc.add_paragraph("Ms. Angela Rivera")
    doc.add_paragraph("Director of Operations")
    doc.add_paragraph("Pacific Commerce Group")
    doc.add_paragraph("789 Harbor View Blvd")
    doc.add_paragraph("Oakland, CA 94607")

    doc.add_paragraph()

    # Salutation
    doc.add_paragraph("Dear Ms. Rivera,")

    doc.add_paragraph()

    # Body paragraphs
    body1 = doc.add_paragraph(
        "Thank you for your continued partnership with Nexus Technology Solutions. "
        "We are pleased to inform you that the integration project has reached its final milestone "
        "ahead of the scheduled completion date. Our engineering team has worked diligently to ensure "
        "that all deliverables meet the performance benchmarks outlined in our agreement."
    )

    doc.add_paragraph()

    body2 = doc.add_paragraph(
        "We would like to schedule a briefing session at our headquarters to present the final results "
        "and discuss the roadmap for Phase 2 implementation. Please let us know your availability "
        "during the week of March 24-28, 2025. A full technical report will be provided prior to "
        "the meeting for your review."
    )

    doc.add_paragraph()

    body3 = doc.add_paragraph(
        "If you require any additional information or have questions in the meantime, please do not "
        "hesitate to contact our project lead, David Okafor, at d.okafor@nexustech.com or by phone "
        "at (408) 555-0192."
    )

    doc.add_paragraph()

    # Closing
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()

    sign_name = doc.add_paragraph()
    run = sign_name.add_run("Katherine Leung")
    run.bold = True

    doc.add_paragraph("Vice President, Client Relations")
    doc.add_paragraph("Nexus Technology Solutions")

    # Company address placeholder (to be filled with AutoCorrect expansion)
    # NOTE: The AutoCorrect entry for 'addr' does NOT exist in initial state
    doc.add_paragraph("[Company Address]")
    doc.add_paragraph("Tel: (408) 555-0100  |  Email: info@nexustech.com")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open the template in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
