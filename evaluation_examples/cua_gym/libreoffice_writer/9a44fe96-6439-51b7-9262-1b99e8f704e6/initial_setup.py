"""
Initial Setup: Corporate template with Heading 2 paragraphs using default style
Task ID: writer_biz_045
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_045'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Title --
    title = doc.add_heading('Meridian Industries — Q1 2025 Strategic Review', level=1)

    # -- Intro paragraph --
    doc.add_paragraph(
        'This document summarizes the key strategic initiatives, departmental performance, '
        'and forward-looking priorities for Meridian Industries during the first quarter of 2025. '
        'All division heads are expected to review the relevant sections and provide feedback '
        'by March 28, 2025.'
    )

    # -- Heading 2: Revenue Performance --
    doc.add_heading('Revenue Performance', level=2)
    doc.add_paragraph(
        'Total revenue for Q1 reached $14.7 million, representing a 12% increase compared '
        'to the same period last year. The growth was primarily driven by the expansion of '
        'our enterprise software licensing division, which contributed $6.3 million. The '
        'professional services arm generated $4.1 million, while hardware solutions accounted '
        'for the remaining $4.3 million.'
    )
    doc.add_paragraph(
        'Recurring revenue from subscription-based products now represents 43% of total '
        'revenue, up from 37% in Q4 2024. Management expects this share to exceed 50% '
        'by year-end as more clients transition to the SaaS model.'
    )

    # -- Heading 2: Operational Highlights --
    doc.add_heading('Operational Highlights', level=2)
    doc.add_paragraph(
        'The operations team completed the migration of our primary data center to the '
        'AWS cloud infrastructure ahead of schedule. This transition is projected to reduce '
        'annual hosting costs by approximately $820,000 while improving uptime guarantees '
        'from 99.5% to 99.95%.'
    )
    doc.add_paragraph(
        'Employee headcount grew from 312 to 341 during the quarter, with 18 new hires '
        'in engineering, 6 in sales, and 5 in customer success. The voluntary attrition '
        'rate remained low at 3.2%, well below the industry average of 8.5%.'
    )

    # -- Heading 2: Product Development --
    doc.add_heading('Product Development', level=2)
    doc.add_paragraph(
        'Version 4.2 of the Meridian Analytics Platform was released on February 15, '
        'incorporating real-time dashboard capabilities and enhanced API integrations. '
        'Early adoption metrics show a 28% increase in daily active users within the '
        'first two weeks of launch.'
    )
    doc.add_paragraph(
        'The product roadmap for Q2 includes the launch of a mobile companion app, '
        'AI-powered anomaly detection features, and a redesigned onboarding workflow '
        'aimed at reducing time-to-value for new enterprise customers.'
    )

    # -- Heading 2: Client Acquisition --
    doc.add_heading('Client Acquisition', level=2)
    doc.add_paragraph(
        'Meridian signed 14 new enterprise contracts during Q1, including notable wins '
        'with Havenport Financial Group, NovaTech Manufacturing, and Cascade Health Systems. '
        'The average contract value increased to $285,000, up from $240,000 in the prior quarter.'
    )
    doc.add_paragraph(
        'The sales pipeline for Q2 currently stands at $9.8 million across 47 qualified '
        'opportunities. The team is focused on converting at least 30% of these to closed '
        'deals by the end of June.'
    )

    # -- Heading 2: Risk Management --
    doc.add_heading('Risk Management', level=2)
    doc.add_paragraph(
        'Key risk areas identified during the quarter include supply chain volatility '
        'affecting hardware margins, increased competition in the mid-market analytics '
        'segment, and potential regulatory changes in the European data privacy landscape. '
        'The risk committee has proposed mitigation strategies for each, to be reviewed '
        'at the April board meeting.'
    )

    # -- Closing paragraph --
    doc.add_paragraph(
        'The executive leadership team remains confident in the trajectory for 2025. '
        'Continued investment in product innovation, talent development, and customer '
        'success will be central to achieving our annual targets.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
