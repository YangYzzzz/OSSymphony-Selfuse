"""
Reward Script: Verify custom TOC entry structure with colon separator
Task ID: writer_mt_083
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): At least one TOC entry has 'N: Text' colon format
  Component 2 (0.3): ALL TOC entries consistently use colon format
  Component 3 (0.3): Body headings remain intact after TOC modification
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_083'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Identify TOC entries: paragraphs between "Table of Contents" header and the first Heading paragraph
    # TOC entries are Normal-style paragraphs that match pattern: number text tab page_number
    toc_entries = []
    in_toc = False
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ''

        if text == 'Table of Contents':
            in_toc = True
            continue

        if in_toc:
            # Stop at first heading (body content starts)
            if 'Heading' in style_name:
                break
            # Skip empty paragraphs
            if not text:
                continue
            # This should be a TOC entry
            toc_entries.append(text)

    print(f"INFO: Found {len(toc_entries)} TOC entries")
    for entry in toc_entries:
        print(f"  TOC: {entry!r}")

    if len(toc_entries) == 0:
        print("FAIL: No TOC entries found")
        print("REWARD: 0.0")
        return 0.0

    # Expected TOC entries from the task: entries should match pattern "N: Text\tPage"
    # e.g., "1: Introduction\t3", "1.1: Background\t4"
    # The colon pattern: number(s) followed by ": " then text
    colon_pattern = re.compile(r'^\d+(\.\d+)*:\s+\S')

    # Component 1: At least one TOC entry has the colon format (0.4 points)
    # This is the core task change - adding ": " between number and text
    try:
        colon_count = sum(1 for entry in toc_entries if colon_pattern.search(entry))
        if colon_count > 0:
            print(f"PASS: Component 1 — {colon_count}/{len(toc_entries)} TOC entries have colon format (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — No TOC entries have colon format. Sample: {toc_entries[0]!r}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: ALL TOC entries consistently use the colon format (0.3 points)
    # Verifies consistency across all levels (1, 1.1, 1.2, 2, etc.)
    try:
        if colon_count == len(toc_entries):
            print(f"PASS: Component 2 — All {len(toc_entries)} TOC entries consistently use colon format (0.3 pts)")
            total_score += 0.3
        else:
            non_colon = [e for e in toc_entries if not colon_pattern.search(e)]
            print(f"FAIL: Component 2 — {len(non_colon)} entries lack colon format: {non_colon[:3]}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Body headings remain intact (0.3 points)
    # The task only changes the TOC structure, not the actual headings
    # Expected headings from the document
    expected_headings = [
        ('Heading 1', 'Introduction'),
        ('Heading 2', 'Background'),
        ('Heading 2', 'Scope'),
        ('Heading 1', 'Methods'),
        ('Heading 2', 'Data Collection'),
        ('Heading 2', 'Analysis Framework'),
        ('Heading 1', 'Results'),
        ('Heading 2', 'Primary Findings'),
        ('Heading 2', 'Secondary Observations'),
        ('Heading 1', 'Discussion'),
        ('Heading 1', 'Conclusion'),
    ]
    try:
        actual_headings = []
        for para in doc.paragraphs:
            if para.style and 'Heading' in para.style.name:
                actual_headings.append((para.style.name, para.text.strip()))

        # Check that all expected headings are present
        matched = 0
        for exp_style, exp_text in expected_headings:
            for act_style, act_text in actual_headings:
                if exp_style == act_style and exp_text == act_text:
                    matched += 1
                    break

        if matched == len(expected_headings):
            print(f"PASS: Component 3 — All {matched} body headings intact (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — Only {matched}/{len(expected_headings)} headings match. "
                  f"Actual: {actual_headings[:5]}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
