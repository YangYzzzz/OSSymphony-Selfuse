"""
Initial Setup: Create a document with chapter numbering and a default TOC
Task ID: writer_mt_083
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_083'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_outline_numbering(doc):
    """Add multilevel list numbering for headings (chapter numbering)."""
    # Create abstractNum and num elements for chapter numbering
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Define abstract numbering with chapter numbering pattern
    abstract_num_xml = f'''
    <w:abstractNum w:abstractNumId="100" {nsdecls('w')}>
        <w:multiLevelType w:val="multilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading1"/>
            <w:lvlText w:val="%1"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="432" w:hanging="432"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading2"/>
            <w:lvlText w:val="%1.%2"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="576" w:hanging="576"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading3"/>
            <w:lvlText w:val="%1.%2.%3"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="720"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    '''
    abstract_num = parse_xml(abstract_num_xml)
    numbering_elm.append(abstract_num)

    num_xml = f'''
    <w:num w:numId="100" {nsdecls('w')}>
        <w:abstractNumId w:val="100"/>
    </w:num>
    '''
    num_elm = parse_xml(num_xml)
    numbering_elm.append(num_elm)


def apply_numbering_to_heading(para, num_id=100, ilvl=0):
    """Apply numbering to a heading paragraph."""
    pPr = para._element.get_or_add_pPr()
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="{ilvl}"/>'
        f'  <w:numId w:val="{num_id}"/>'
        f'</w:numPr>'
    )
    pPr.append(numPr)


def add_toc_field(doc):
    """Add a Table of Contents field with default entry structure (E# E T#)."""
    # Add a TOC title
    toc_title = doc.add_paragraph()
    toc_title.style = doc.styles['Normal']
    run = toc_title.add_run('Table of Contents')
    run.bold = True
    run.font.size = Pt(16)
    toc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_title.paragraph_format.space_after = Pt(12)

    # Add TOC field - default structure shows "E# E T#"
    # In OOXML, TOC is a field code
    toc_para = doc.add_paragraph()

    # Begin field
    r1 = toc_para.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    # Field instruction - TOC with default switches
    r2 = toc_para.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._element.append(instr)

    # Separate
    r3 = toc_para.add_run()
    fld_sep = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r3._element.append(fld_sep)

    # Placeholder entries showing default format: "E# E T#"
    # (these are cached display entries; LibreOffice will regenerate on update)
    default_entries = [
        ('1 Introduction', '3', 'TOC 1'),
        ('1.1 Background', '4', 'TOC 2'),
        ('1.2 Scope', '5', 'TOC 2'),
        ('2 Methods', '8', 'TOC 1'),
        ('2.1 Data Collection', '9', 'TOC 2'),
        ('2.2 Analysis Framework', '11', 'TOC 2'),
        ('3 Results', '14', 'TOC 1'),
        ('3.1 Primary Findings', '15', 'TOC 2'),
        ('3.2 Secondary Observations', '17', 'TOC 2'),
        ('4 Discussion', '20', 'TOC 1'),
        ('5 Conclusion', '24', 'TOC 1'),
    ]

    for entry_text, page_num, style_name in default_entries:
        entry_para = doc.add_paragraph()
        # Set TOC style via XML
        pPr = entry_para._element.get_or_add_pPr()
        pStyle = pPr.makeelement(qn('w:pStyle'), {qn('w:val'): style_name.replace(' ', '')})
        pPr.append(pStyle)

        # Add tab stop with dot leader for page number
        tabs_elm = parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'  <w:tab w:val="right" w:leader="dot" w:pos="9072"/>'
            f'</w:tabs>'
        )
        pPr.append(tabs_elm)

        entry_run = entry_para.add_run(entry_text)
        tab_run = entry_para.add_run('\t')
        page_run = entry_para.add_run(page_num)

    # End field
    end_para = doc.add_paragraph()
    r_end = end_para.add_run()
    fld_end = r_end._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r_end._element.append(fld_end)

    # Add page break after TOC
    doc.add_page_break()


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Configure heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Arial'
        if level == 1:
            h_style.font.size = Pt(16)
        elif level == 2:
            h_style.font.size = Pt(14)
        else:
            h_style.font.size = Pt(12)

    # Ensure numbering part exists by adding a dummy list item then removing it
    dummy = doc.add_paragraph('dummy', style='List Number')
    dummy._element.getparent().remove(dummy._element)

    # Add outline numbering
    add_outline_numbering(doc)

    # Title page
    title = doc.add_heading('Numbered Report', level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Annual Research Analysis 2025')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(24)
    run = author.add_run('Prepared by Dr. Sarah Chen\nResearch Division')
    run.font.size = Pt(11)

    doc.add_page_break()

    # Add TOC with default entry structure
    add_toc_field(doc)

    # Chapter 1: Introduction
    h1 = doc.add_heading('Introduction', level=1)
    apply_numbering_to_heading(h1, ilvl=0)

    doc.add_paragraph(
        'This report presents a comprehensive analysis of the research findings '
        'gathered during the fiscal year 2025. The study encompasses multiple '
        'departments and evaluates key performance indicators across various '
        'operational domains.'
    )
    doc.add_paragraph(
        'The methodology employed in this investigation follows established '
        'protocols from the International Research Standards Organization (IRSO), '
        'ensuring reproducibility and statistical rigor throughout all phases.'
    )

    # 1.1 Background
    h11 = doc.add_heading('Background', level=2)
    apply_numbering_to_heading(h11, ilvl=1)

    doc.add_paragraph(
        'The organizational restructuring initiated in Q3 2024 prompted a '
        'thorough review of existing research methodologies. Previous annual '
        'reports indicated a 23% improvement in data collection efficiency, '
        'which served as the baseline for this year\'s targets.'
    )
    doc.add_paragraph(
        'Historical data from 2020-2024 demonstrates a consistent upward '
        'trend in research output quality, with peer-reviewed publication '
        'rates increasing from 34% to 67% over the five-year period.'
    )

    # 1.2 Scope
    h12 = doc.add_heading('Scope', level=2)
    apply_numbering_to_heading(h12, ilvl=1)

    doc.add_paragraph(
        'This analysis covers all research activities conducted between '
        'January 1, 2025 and December 31, 2025. The scope includes '
        'laboratory experiments, field studies, computational modeling, '
        'and cross-departmental collaborative projects.'
    )

    doc.add_page_break()

    # Chapter 2: Methods
    h2 = doc.add_heading('Methods', level=1)
    apply_numbering_to_heading(h2, ilvl=0)

    doc.add_paragraph(
        'The research methodology was designed to maximize data integrity '
        'while maintaining practical feasibility across all participating '
        'departments. A mixed-methods approach was adopted, combining '
        'quantitative analysis with qualitative assessments.'
    )

    # 2.1 Data Collection
    h21 = doc.add_heading('Data Collection', level=2)
    apply_numbering_to_heading(h21, ilvl=1)

    doc.add_paragraph(
        'Data was collected through three primary channels: automated sensor '
        'networks deployed across 14 research facilities, structured interviews '
        'with 237 research personnel, and archival analysis of 1,842 published '
        'papers from the preceding decade.'
    )
    doc.add_paragraph(
        'Quality assurance protocols required dual verification of all data '
        'points, with a maximum acceptable error margin of 2.5%. Outliers '
        'exceeding three standard deviations were flagged for manual review.'
    )

    # 2.2 Analysis Framework
    h22 = doc.add_heading('Analysis Framework', level=2)
    apply_numbering_to_heading(h22, ilvl=1)

    doc.add_paragraph(
        'Statistical analysis was performed using a combination of parametric '
        'and non-parametric tests. The primary framework utilized hierarchical '
        'linear modeling (HLM) to account for nested data structures within '
        'departmental clusters.'
    )

    doc.add_page_break()

    # Chapter 3: Results
    h3 = doc.add_heading('Results', level=1)
    apply_numbering_to_heading(h3, ilvl=0)

    doc.add_paragraph(
        'The analysis yielded statistically significant results across '
        'multiple domains. Key findings are organized by research priority '
        'level and departmental impact.'
    )

    # 3.1 Primary Findings
    h31 = doc.add_heading('Primary Findings', level=2)
    apply_numbering_to_heading(h31, ilvl=1)

    doc.add_paragraph(
        'Revenue from research grants increased by 31% compared to the '
        'previous fiscal year, totaling $4.2 million. The Engineering '
        'department contributed 42% of all grant funding, followed by '
        'Biomedical Sciences at 28% and Environmental Studies at 18%.'
    )
    doc.add_paragraph(
        'Patent applications filed during the reporting period reached '
        'an all-time high of 47, representing a 56% increase over 2024. '
        'Of these, 23 have already received preliminary approval from '
        'the relevant patent offices.'
    )

    # 3.2 Secondary Observations
    h32 = doc.add_heading('Secondary Observations', level=2)
    apply_numbering_to_heading(h32, ilvl=1)

    doc.add_paragraph(
        'Cross-departmental collaboration metrics showed that joint '
        'projects produced 2.3 times more citations than single-department '
        'publications. This finding supports the institutional strategy '
        'of encouraging interdisciplinary research initiatives.'
    )

    doc.add_page_break()

    # Chapter 4: Discussion
    h4 = doc.add_heading('Discussion', level=1)
    apply_numbering_to_heading(h4, ilvl=0)

    doc.add_paragraph(
        'The results of this annual analysis confirm several hypotheses '
        'proposed in the 2024 strategic planning document. The significant '
        'increase in grant revenue and patent applications suggests that '
        'the organizational restructuring has had a measurable positive '
        'impact on research productivity.'
    )
    doc.add_paragraph(
        'However, it is important to note that the correlation between '
        'restructuring and improved outcomes does not necessarily imply '
        'causation. External factors, including increased federal research '
        'funding and favorable market conditions for technology transfer, '
        'may have contributed to the observed improvements.'
    )
    doc.add_paragraph(
        'Future research should focus on isolating the specific variables '
        'responsible for performance gains, with particular attention to '
        'the role of mentorship programs and early-career researcher support '
        'initiatives launched in mid-2024.'
    )

    doc.add_page_break()

    # Chapter 5: Conclusion
    h5 = doc.add_heading('Conclusion', level=1)
    apply_numbering_to_heading(h5, ilvl=0)

    doc.add_paragraph(
        'This report demonstrates that the research division has achieved '
        'substantial progress across all key performance indicators during '
        'the 2025 fiscal year. The combination of increased grant funding, '
        'higher publication rates, and record patent applications positions '
        'the organization favorably for continued growth.'
    )
    doc.add_paragraph(
        'Based on these findings, we recommend maintaining the current '
        'organizational structure while investing in expanded data analytics '
        'capabilities to further optimize research operations in 2026 and '
        'beyond.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
