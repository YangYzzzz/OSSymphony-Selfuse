"""
Initial Setup: Print preparation checklist task
Task ID: osworld_multi_apps_writer_gimp_073
Domain: libreoffice_writer + gimp (multi-app)

Creates:
  - /home/user/print_prep.docx: print preparation checklist document
  - /home/user/Desktop/poster.png: RGB poster image at 72 DPI, 595x842 px
Launches:
  - LibreOffice Writer with print_prep.docx
"""

import os
import shlex
import subprocess
import time

# ---- python-docx ----
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ---- Pillow ----
from PIL import Image, ImageDraw, ImageFont
import struct

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_writer_gimp_073'
DOCX_PATH = f'{WORKDIR}/print_prep.docx'
POSTER_PATH = f'{WORKDIR}/Desktop/poster.png'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_png_dpi(path: str, dpi: int):
    """Embed DPI metadata into PNG file (pHYs chunk, unit=meters)."""
    ppm = int(dpi / 0.0254)  # pixels per meter
    with open(path, 'rb') as f:
        data = f.read()

    # Build pHYs chunk: ppuX (4 bytes), ppuY (4 bytes), unit (1 byte = 1 for meters)
    phys_data = struct.pack('>IIB', ppm, ppm, 1)
    import zlib
    crc = zlib.crc32(b'pHYs' + phys_data) & 0xffffffff
    phys_chunk = struct.pack('>I', 9) + b'pHYs' + phys_data + struct.pack('>I', crc)

    # Insert pHYs chunk after PNG signature (8 bytes) + IHDR chunk (4+4+13+4 = 25 bytes)
    sig = data[:8]
    ihdr = data[8:33]  # 4 len + 4 'IHDR' + 13 data + 4 CRC
    rest = data[33:]

    with open(path, 'wb') as f:
        f.write(sig + ihdr + phys_chunk + rest)


def create_print_prep_docx():
    """Create the print preparation checklist document."""
    doc = Document()

    # Title
    title = doc.add_heading('Print Preparation Checklist', level=0)
    title.paragraph_format.space_after = Pt(12)

    # Introduction paragraph
    intro = doc.add_paragraph(
        'Before sending your poster to the print shop, please follow this checklist to ensure '
        'your file meets professional printing standards. Apply each step carefully to the file.'
    )
    intro.paragraph_format.space_after = Pt(10)

    # Section heading
    doc.add_heading('Pre-Press Checklist Steps', level=1)

    # Step 1: CMYK Conversion
    step1_heading = doc.add_paragraph()
    run = step1_heading.add_run('Step 1: Convert to CMYK Color Mode (Simulated)')
    run.bold = True
    run.font.size = Pt(12)
    step1_heading.paragraph_format.space_before = Pt(8)
    step1_heading.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        'Professional printing presses use CMYK (Cyan, Magenta, Yellow, Black) ink. '
        'RGB files must be converted to CMYK for accurate color reproduction. '
        'In GIMP, use the CMYK separation plugin (Filters > Colors > Separate+) to simulate '
        'CMYK output, then flatten and export. This step ensures colors print as expected.'
    )

    doc.add_paragraph('Actions required:', style='List Bullet')
    doc.add_paragraph('Open poster.png in GIMP', style='List Bullet')
    doc.add_paragraph('Apply CMYK color simulation via desaturation/channel manipulation', style='List Bullet')
    doc.add_paragraph('Flatten image to merge all layers', style='List Bullet')

    # Step 2: Resolution
    step2_heading = doc.add_paragraph()
    run = step2_heading.add_run('Step 2: Set Resolution to 300 DPI')
    run.bold = True
    run.font.size = Pt(12)
    step2_heading.paragraph_format.space_before = Pt(8)
    step2_heading.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        'Print quality requires a minimum resolution of 300 DPI (dots per inch). '
        'The current poster.png is at 72 DPI (screen resolution), which will appear '
        'blurry when printed. Resample the image to 300 DPI to ensure sharp print output.'
    )

    doc.add_paragraph('Actions required:', style='List Bullet')
    doc.add_paragraph(
        'In GIMP: Image > Scale Image, set X and Y resolution to 300 pixels/in',
        style='List Bullet'
    )
    doc.add_paragraph('Use cubic interpolation for best quality', style='List Bullet')
    doc.add_paragraph('Confirm the canvas size increases proportionally', style='List Bullet')

    # Step 3: Bleed
    step3_heading = doc.add_paragraph()
    run = step3_heading.add_run('Step 3: Add 3mm Bleed Margins')
    run.bold = True
    run.font.size = Pt(12)
    step3_heading.paragraph_format.space_before = Pt(8)
    step3_heading.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        'Bleed is extra image area beyond the final trim edge that prevents white borders '
        'when the printer cuts the paper. A standard 3mm bleed must be added on all four sides '
        'of the poster. Extend the canvas by 3mm on each side and fill with background color '
        'or extend existing artwork.'
    )

    doc.add_paragraph('Actions required:', style='List Bullet')
    doc.add_paragraph(
        'In GIMP: Image > Canvas Size, add 6mm to width and height (3mm each side)',
        style='List Bullet'
    )
    doc.add_paragraph('Center the existing content (offset by 3mm)', style='List Bullet')
    doc.add_paragraph('Flatten and fill the bleed area with appropriate background', style='List Bullet')

    # Export section
    doc.add_heading('Final Export', level=1)
    export_para = doc.add_paragraph(
        'Once all three steps above are complete, export the file as a PDF for print delivery:'
    )
    export_para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'File > Export As > poster_print_ready.pdf (save to Desktop)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'In PDF export options: select "Flatten image" and ensure compatibility mode is set',
        style='List Bullet'
    )

    # Notes section
    doc.add_heading('Important Notes', level=2)
    notes = doc.add_paragraph(
        'The output file must be saved as poster_print_ready.pdf on the Desktop. '
        'Verify the file exists before submitting to the print shop. '
        'The final PDF should reflect: simulated CMYK color profile, 300 DPI resolution, '
        'and 3mm bleed on all sides.'
    )
    notes.paragraph_format.space_before = Pt(6)

    doc.save(DOCX_PATH)
    print(f'Created: {DOCX_PATH}')


def create_poster_png():
    """Create a realistic A4 poster image at 72 DPI, 595x842 px (RGB)."""
    width, height = 595, 842  # A4 at 72 DPI
    img = Image.new('RGB', (width, height), color=(240, 235, 220))
    draw = ImageDraw.Draw(img)

    # Background gradient effect using bands
    for y in range(height):
        factor = y / height
        r = int(240 - factor * 20)
        g = int(235 - factor * 30)
        b = int(220 - factor * 10)
        draw.line([(0, y), (width, y)], fill=(r, g, b))

    # Header bar
    draw.rectangle([0, 0, width, 100], fill=(52, 73, 94))

    # Title text (using default font since truetype may not be available)
    try:
        title_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 36)
        subtitle_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 18)
        body_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 14)
        small_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 12)
    except Exception:
        title_font = ImageFont.load_default()
        subtitle_font = title_font
        body_font = title_font
        small_font = title_font

    # Title
    draw.text((30, 25), 'ANNUAL DESIGN EXHIBITION', font=title_font, fill=(255, 255, 255))
    draw.text((30, 68), '2025 Creative Arts Showcase', font=subtitle_font, fill=(189, 195, 199))

    # Main decorative rectangle
    draw.rectangle([30, 120, width - 30, 360], fill=(70, 130, 180), outline=(52, 73, 94), width=3)

    # Inner decorative elements
    draw.ellipse([width // 2 - 80, 150, width // 2 + 80, 330], fill=(255, 200, 50), outline=(200, 150, 30), width=2)

    # Featured text
    draw.text((width // 2 - 60, 220), 'FEATURED', font=subtitle_font, fill=(52, 73, 94))
    draw.text((width // 2 - 70, 250), 'COLLECTION', font=subtitle_font, fill=(52, 73, 94))

    # Section 1
    draw.rectangle([30, 380, 270, 550], fill=(231, 76, 60), outline=(192, 57, 43), width=2)
    draw.text((50, 395), 'Photography', font=subtitle_font, fill=(255, 255, 255))
    draw.text((50, 425), '42 works by local artists', font=body_font, fill=(255, 235, 230))
    draw.text((50, 450), 'Gallery Hall A - Level 2', font=body_font, fill=(255, 235, 230))
    draw.text((50, 480), 'Opening: March 15, 2025', font=small_font, fill=(255, 200, 190))
    draw.text((50, 500), 'Daily: 10:00 - 18:00', font=small_font, fill=(255, 200, 190))
    draw.text((50, 520), 'Admission: Free', font=small_font, fill=(255, 220, 210))

    # Section 2
    draw.rectangle([300, 380, width - 30, 550], fill=(39, 174, 96), outline=(27, 133, 70), width=2)
    draw.text((320, 395), 'Digital Art', font=subtitle_font, fill=(255, 255, 255))
    draw.text((320, 425), '28 interactive installations', font=body_font, fill=(220, 255, 235))
    draw.text((320, 450), 'Gallery Hall B - Level 1', font=body_font, fill=(220, 255, 235))
    draw.text((320, 480), 'Opening: March 16, 2025', font=small_font, fill=(180, 240, 200))
    draw.text((320, 500), 'Daily: 11:00 - 20:00', font=small_font, fill=(180, 240, 200))
    draw.text((320, 520), 'Admission: $5 suggested', font=small_font, fill=(200, 255, 220))

    # Section 3 - workshops
    draw.rectangle([30, 570, width - 30, 680], fill=(142, 68, 173), outline=(113, 54, 138), width=2)
    draw.text((50, 585), 'Workshops & Talks', font=subtitle_font, fill=(255, 255, 255))
    draw.text((50, 615), 'Saturday March 22: Print Techniques (10:00-12:00)', font=body_font, fill=(230, 200, 255))
    draw.text((50, 638), 'Sunday March 23: Digital Color Theory (14:00-16:00)', font=body_font, fill=(230, 200, 255))
    draw.text((50, 661), 'Register at info@designexhibit.com or call +1-555-0147', font=small_font, fill=(210, 180, 240))

    # Footer
    draw.rectangle([0, 750, width, 842], fill=(44, 62, 80))
    draw.text((30, 760), 'City Arts Center  |  123 Exhibition Boulevard  |  www.cityarts.example.com', font=small_font, fill=(149, 165, 166))
    draw.text((30, 790), 'Follow us: @CityArtsCenter  |  Presented by the Municipal Arts Council', font=small_font, fill=(127, 140, 141))
    draw.text((30, 815), 'Supported by: Creative Fund Grant 2025', font=small_font, fill=(100, 120, 130))

    # Ensure Desktop directory exists
    desktop_dir = os.path.dirname(POSTER_PATH)
    os.makedirs(desktop_dir, exist_ok=True)

    # Save as PNG (will be 72 DPI by default in Pillow)
    img.save(POSTER_PATH, format='PNG', dpi=(72, 72))

    # Embed 72 DPI in pHYs chunk
    try:
        set_png_dpi(POSTER_PATH, 72)
    except Exception as e:
        print(f'Warning: could not embed DPI: {e}')

    print(f'Created: {POSTER_PATH} (595x842, RGB, 72 DPI)')


def create_initial():
    create_print_prep_docx()
    create_poster_png()

    # GUI-ready startup: open print_prep.docx in LibreOffice Writer first
    launch_gui(f'libreoffice --writer "{DOCX_PATH}"', delay_sec=3.0)
    # Also open GIMP with the poster image
    launch_gui(f'gimp "{POSTER_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer and GIMP with DISPLAY=:0')


create_initial()
