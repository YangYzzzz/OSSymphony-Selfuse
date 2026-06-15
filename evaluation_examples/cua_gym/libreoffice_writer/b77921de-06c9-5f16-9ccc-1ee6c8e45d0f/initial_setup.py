"""
Initial Setup: Configure Heading 1 style in LibreOffice Writer
Task ID: writer_tech_080
Domain: libreoffice_writer

Creates a technical specification document with default Heading 1 styling
(Liberation Sans 24pt bold). The document contains realistic tech content
with multiple Heading 1 and Heading 2 sections.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_080'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Title --
    title = doc.add_heading("CloudScale Platform — Technical Architecture", level=0)

    # -- Section 1: Heading 1 --
    doc.add_heading("System Overview", level=1)
    doc.add_paragraph(
        "CloudScale is a distributed cloud infrastructure platform designed to handle "
        "enterprise-grade workloads with automatic horizontal scaling. The platform "
        "processes over 2.4 million API requests per minute across 14 global regions, "
        "maintaining a 99.97% uptime SLA for all production services."
    )
    doc.add_paragraph(
        "Built on a microservices architecture, the platform uses event-driven "
        "communication via Apache Kafka for inter-service messaging and gRPC for "
        "synchronous operations requiring sub-10ms latency."
    )

    # -- Section 1.1: Heading 2 --
    doc.add_heading("Core Components", level=2)
    doc.add_paragraph(
        "The platform consists of five primary subsystems: the API Gateway (Kong-based), "
        "the Compute Orchestrator (Kubernetes 1.28), the Storage Layer (Ceph + S3-compatible "
        "interface), the Networking Fabric (Cilium eBPF), and the Observability Stack "
        "(Prometheus, Grafana, Jaeger)."
    )

    # -- Section 2: Heading 1 --
    doc.add_heading("Authentication and Authorization", level=1)
    doc.add_paragraph(
        "All client requests are authenticated using OAuth 2.0 with PKCE flow. Service-to-service "
        "communication relies on mTLS certificates issued by an internal Vault PKI backend. "
        "Authorization is enforced through Open Policy Agent (OPA) with Rego policies evaluated "
        "at the API Gateway level before requests reach backend services."
    )

    doc.add_heading("Token Management", level=2)
    doc.add_paragraph(
        "Access tokens have a 15-minute TTL with refresh tokens valid for 7 days. Token "
        "rotation is handled by the Identity Service, which maintains a Redis-backed session "
        "store with geographic affinity routing. Compromised tokens can be revoked within "
        "30 seconds through the global revocation list propagated via Kafka."
    )

    # -- Section 3: Heading 1 --
    doc.add_heading("Data Pipeline Architecture", level=1)
    doc.add_paragraph(
        "The real-time data pipeline ingests events from 340+ microservices through a "
        "centralized Kafka cluster with 96 brokers. Events are processed using Apache Flink "
        "for stream analytics and Apache Spark for batch aggregation jobs. The pipeline handles "
        "approximately 1.8 TB of raw event data daily."
    )

    doc.add_heading("Event Schema Registry", level=2)
    doc.add_paragraph(
        "All event schemas are managed through Confluent Schema Registry with Avro serialization. "
        "Schema evolution follows backward-compatible rules enforced at the registry level. "
        "Breaking changes require a new topic version and a 30-day migration window with "
        "dual-publishing to both old and new topics."
    )

    # -- Section 4: Heading 1 --
    doc.add_heading("Deployment Strategy", level=1)
    doc.add_paragraph(
        "Production deployments follow a blue-green strategy with canary validation. Each "
        "release candidate is first deployed to a canary cluster serving 5% of traffic for "
        "a minimum of 2 hours. Automated rollback triggers if error rates exceed 0.1% or "
        "p99 latency increases by more than 15% compared to the baseline."
    )
    doc.add_paragraph(
        "Infrastructure changes are managed through Terraform with state stored in an "
        "encrypted S3 backend. All changes require peer review via GitOps workflow and "
        "must pass automated policy checks before being applied to production environments."
    )

    # -- Section 5: Heading 1 --
    doc.add_heading("Disaster Recovery", level=1)
    doc.add_paragraph(
        "The platform maintains active-passive replication across three geographic zones "
        "with an RPO of 30 seconds and RTO of 5 minutes. Database backups are taken every "
        "6 hours using pg_dump with WAL archiving for point-in-time recovery. Cross-region "
        "failover is orchestrated through AWS Route 53 health checks with automatic DNS "
        "cutover when primary region health drops below 95%."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
