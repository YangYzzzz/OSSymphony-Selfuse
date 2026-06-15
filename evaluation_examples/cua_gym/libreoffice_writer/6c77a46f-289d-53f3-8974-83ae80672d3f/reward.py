"""
FINAL REWARD SCRIPT - SUCCESS
Task: Convert the 4 in "SO4" to subscript in the reaction note.
Generated: 2025-10-14 05:43:37
Status: success
Model: azure-o3
Total Steps: 1
"""

from docx import Document
import os

def find_subscript_SO4(doc):
    """Return True if there is any occurrence of the text pattern
    S O (subscript 4) in the document. We confirm that:
      1. A run containing the single character '4' has the font.subscript flag set
      2. The immediately preceding visible characters in the same paragraph are
         exactly the letters 'SO'.
    """
    for para_idx, paragraph in enumerate(doc.paragraphs):
        runs = paragraph.runs
        for run_idx, run in enumerate(runs):
            # python-docx represents the subscript attribute as a boolean or None
            subscript_flag = getattr(run.font, "subscript", None)
            # We only care about a run that is just the character "4" and is subscript
            if run.text == "4" and subscript_flag:
                # Build the paragraph text up to but **excluding** this run
                preceding_text = "".join(r.text for r in runs[:run_idx])
                if preceding_text.endswith("SO"):
                    print(f"✓ Found subscript '4' in 'SO4' (paragraph {para_idx})")
                    return True
    print("✗ Could not find subscript '4' in 'SO4'")
    return False

def has_plain_SO4(doc):
    """Detect if any plain (non-subscript) 'SO4' remains in the document.
    We look for a run that contains the digit 4 **without** the subscript flag
    and that is preceded by the letters 'SO'. If such an occurrence exists, it
    means the task was not fully completed (both old and new forms coexist)."""
    for para_idx, paragraph in enumerate(doc.paragraphs):
        runs = paragraph.runs
        for run_idx, run in enumerate(runs):
            subscript_flag = getattr(run.font, "subscript", None)
            # Any non-subscript run that includes '4' is suspect
            if "4" in run.text and not subscript_flag:
                # Split at the first 4 to see what comes immediately before it
                before_four = run.text.split("4")[0]
                preceding_text = "".join(r.text for r in runs[:run_idx]) + before_four
                if preceding_text.endswith("SO"):
                    print(
                        f"Found plain 'SO4' (non-subscript 4) still present in paragraph {para_idx}"
                    )
                    return True
    return False

def verify_task(file_path):
    """Main verification function for the task.

    Scoring rubric (progressive):
      • +0.7  if at least one correctly formatted 'SO₄' (4 as subscript) is found
      • +0.3  bonus if **no** plain 'SO4' remains, ensuring the conversion is clean
        (total possible = 1.0).
    """
    max_score = 1.0
    score = 0.0

    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Error opening document: {e}")
        return 0.0

    # 1) Verify the 4 is in subscript format for at least one occurrence of SO4
    if find_subscript_SO4(doc):
        score += 0.7
        # 2) Ensure the original plain SO4 no longer exists (clean replacement)
        if not has_plain_SO4(doc):
            score += 0.3
        else:
            print("Non-subscript 'SO4' still present – awarding partial credit.")
    else:
        print("Required subscript formatting not detected – no credit awarded.")

    # Cap score at 1.0
    final_score = min(score, max_score)
    print(f"REWARD: {final_score}")
    return final_score

# ----- Execute verification when script is run directly -----
if __name__ == "__main__":
    FILE_PATH = "/home/user/convert_the_4_in_so4_to_subscript_in_the_reaction_note.docx"
    verify_task(FILE_PATH)

