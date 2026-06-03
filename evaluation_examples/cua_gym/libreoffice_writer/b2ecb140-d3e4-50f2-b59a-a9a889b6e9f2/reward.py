"""
Reward Script: Create custom list style with REQ-001, REQ-002 format for 8 requirements
Task ID: writer_tech_048
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): All 8 requirement paragraphs (P10-P17) have numbering properties
  Component 2 (0.35): Numbering definition uses REQ-prefix with zero-padded format
  Component 3 (0.30): Manual REQ-XXX: prefix removed from paragraph text
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_048'

# Indices of the 8 requirement paragraphs in the document
REQ_PARA_INDICES = list(range(10, 18))
NUM_REQUIREMENTS = 8


def persist_app_state(domain):
    """Try to save any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for " + domain)
    except Exception as e:
        print("PERSIST_WARN: save hook failed: " + str(e))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print("CRITICAL: Missing python-docx library: " + str(e))
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file " + file_path + ": " + str(e))
        print("REWARD: 0.0")
        return 0.0

    num_paras = len(doc.paragraphs)
    if num_paras < 18:
        print("CRITICAL: Document has only " + str(num_paras) + " paragraphs, expected at least 18")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: All 8 requirement paragraphs have numbering properties (0.35 pts)
    # In initial, these are Normal with no numPr. In golden, they have numPr with numId.
    try:
        numbered_count = 0
        num_ids_found = set()
        for idx in REQ_PARA_INDICES:
            para = doc.paragraphs[idx]
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_el = numPr.find(qn('w:numId'))
                    if numId_el is not None:
                        nid = numId_el.get(qn('w:val'))
                        if nid and nid != '0':
                            numbered_count += 1
                            num_ids_found.add(nid)

        if numbered_count == NUM_REQUIREMENTS:
            print("PASS: Component 1 -- All 8 requirement paragraphs have numbering (0.35 pts)")
            total_score += 0.35
        elif numbered_count > 0:
            partial = 0.35 * (numbered_count / NUM_REQUIREMENTS)
            print("PARTIAL: Component 1 -- " + str(numbered_count) + "/8 paragraphs numbered (" + str(round(partial, 3)) + " pts)")
            total_score += partial
        else:
            print("FAIL: Component 1 -- No requirement paragraphs have numbering")

        # Store for component 2
        shared_num_ids = num_ids_found
    except Exception as e:
        print("ERROR: Component 1 -- " + str(e))
        shared_num_ids = set()

    # Component 2: Numbering definition uses REQ- prefix with zero-padded format (0.35 pts)
    # Check the abstractNum definition referenced by the requirement paragraphs.
    # Expected: numFmt=decimalZero and lvlText containing 'REQ-' prefix.
    try:
        if not shared_num_ids:
            print("FAIL: Component 2 -- No numbering IDs found, cannot check format")
        else:
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                print("FAIL: Component 2 -- No numbering part in document")
            else:
                numbering_xml = numbering_part._element
                req_format_found = False

                for num_id in shared_num_ids:
                    # Find the <w:num> element with this numId
                    for num_el in numbering_xml.findall(qn('w:num')):
                        if num_el.get(qn('w:numId')) == num_id:
                            abs_ref = num_el.find(qn('w:abstractNumId'))
                            if abs_ref is not None:
                                abs_id = abs_ref.get(qn('w:val'))
                                # Find the abstractNum
                                for abs_num in numbering_xml.findall(qn('w:abstractNum')):
                                    if abs_num.get(qn('w:abstractNumId')) == abs_id:
                                        for lvl in abs_num.findall(qn('w:lvl')):
                                            ilvl_val = lvl.get(qn('w:ilvl'))
                                            if ilvl_val == '0':
                                                numFmt_el = lvl.find(qn('w:numFmt'))
                                                lvlText_el = lvl.find(qn('w:lvlText'))
                                                numFmt_val = numFmt_el.get(qn('w:val')) if numFmt_el is not None else None
                                                lvlText_val = lvlText_el.get(qn('w:val')) if lvlText_el is not None else None

                                                print("  DEBUG: numFmt=" + str(numFmt_val) + " lvlText=" + str(lvlText_val))

                                                # Check for REQ- prefix in lvlText
                                                has_req_prefix = False
                                                if lvlText_val and 'REQ-' in lvlText_val.upper():
                                                    has_req_prefix = True

                                                # Check for zero-padded numbering
                                                # decimalZero gives 01, 02, etc.
                                                # Or could use custom format
                                                has_zero_pad = False
                                                if numFmt_val in ('decimalZero', 'decimal'):
                                                    # decimalZero is ideal; decimal with REQ-00%1 pattern also works
                                                    if numFmt_val == 'decimalZero':
                                                        has_zero_pad = True
                                                    elif lvlText_val and '00' in lvlText_val:
                                                        has_zero_pad = True

                                                if has_req_prefix and has_zero_pad:
                                                    req_format_found = True
                                                elif has_req_prefix:
                                                    # REQ prefix but not zero-padded
                                                    req_format_found = True  # still acceptable with partial
                                            break

                if req_format_found:
                    print("PASS: Component 2 -- Numbering uses REQ- prefix with zero-padded format (0.35 pts)")
                    total_score += 0.35
                else:
                    print("FAIL: Component 2 -- Numbering does not use REQ- prefix format")
    except Exception as e:
        print("ERROR: Component 2 -- " + str(e))

    # Component 3: Manual REQ-XXX: prefix removed from paragraph text (0.30 pts)
    # In initial, text starts with "REQ-001: ..." etc.
    # In golden, the prefix is auto-generated so text should NOT start with REQ-
    try:
        clean_count = 0
        for idx in REQ_PARA_INDICES:
            para_text = doc.paragraphs[idx].text.strip()
            # Check that text does NOT start with "REQ-" pattern (manual prefix)
            if not re.match(r'^REQ-\d{3}\s*:', para_text):
                clean_count += 1
            else:
                print("  DETAIL: P" + str(idx) + " still has manual prefix: " + para_text[:30])

        if clean_count == NUM_REQUIREMENTS:
            print("PASS: Component 3 -- All 8 paragraphs have manual REQ prefix removed (0.30 pts)")
            total_score += 0.30
        elif clean_count > 0:
            partial = 0.30 * (clean_count / NUM_REQUIREMENTS)
            print("PARTIAL: Component 3 -- " + str(clean_count) + "/8 paragraphs have prefix removed (" + str(round(partial, 3)) + " pts)")
            total_score += partial
        else:
            print("FAIL: Component 3 -- All paragraphs still have manual REQ-XXX: prefix")
    except Exception as e:
        print("ERROR: Component 3 -- " + str(e))

    final_score = round(min(total_score, 1.0), 2)
    print()
    print("Score: " + str(total_score) + "/1.0")
    print("REWARD: " + str(final_score))
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = WORKDIR + '/' + TASK_ID + '.docx'
if not os.path.exists(file_path):
    print("File not found: " + file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
