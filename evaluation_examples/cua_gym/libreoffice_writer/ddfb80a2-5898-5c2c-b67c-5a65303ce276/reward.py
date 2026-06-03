"""
Reward Script: Company Newsletter in LibreOffice Writer
Task ID: writer_wf_004
Domain: libreoffice_writer
Scoring:
  Component 1: 2-column layout (0.15 points)
  Component 2: Header with 'Nexora Monthly - Issue 14' in bold (0.20 points)
  Component 3: 4 section headings with Heading 2 style (0.30 points)
  Component 4: Horizontal lines between sections (0.20 points)
  Component 5: Centered page numbers in footer (0.15 points)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_wf_004'

EXPECTED_SECTIONS = ['CEO Message', 'New Product Launch', 'Employee Spotlight', 'Upcoming Events']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: 2-column layout (0.15 points)
    # The section must have w:cols with w:num="2"
    try:
        section = doc.sections[0]
        sectPr = section._sectPr
        cols_elem = sectPr.find(qn('w:cols'))
        if cols_elem is not None:
            num_cols = cols_elem.get(qn('w:num'))
            if num_cols == '2':
                print(f"PASS: Component 1 — 2-column layout found (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 1 — Expected 2 columns, found num={num_cols}")
        else:
            print(f"FAIL: Component 1 — No w:cols element found in section properties")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header with 'Nexora Monthly - Issue 14' in bold (0.20 points)
    # Check header text and bold formatting
    try:
        section = doc.sections[0]
        header = section.header
        header_text = ''
        header_bold = False
        if header and header.paragraphs:
            for p in header.paragraphs:
                if p.text.strip():
                    header_text = p.text.strip()
                    # Check if any run with text is bold
                    runs_with_text = [r for r in p.runs if r.text.strip()]
                    if runs_with_text:
                        header_bold = all(r.bold for r in runs_with_text)
                    break

        text_match = 'Nexora Monthly - Issue 14' in header_text
        if text_match and header_bold:
            print(f"PASS: Component 2 — Header text '{header_text}' is bold (0.20 pts)")
            total_score += 0.20
        elif text_match and not header_bold:
            print(f"PARTIAL: Component 2 — Header text matches but not bold (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 2 — Header text='{header_text}', bold={header_bold}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: 4 section headings with Heading 2 style (0.30 points)
    # Each heading is worth 0.075 points
    try:
        heading2_texts = []
        for para in doc.paragraphs:
            if para.style and para.style.name == 'Heading 2':
                heading2_texts.append(para.text.strip())

        found_count = 0
        for expected in EXPECTED_SECTIONS:
            if any(expected.lower() in h.lower() for h in heading2_texts):
                found_count += 1

        if found_count == 4:
            print(f"PASS: Component 3 — All 4 section headings found with Heading 2 style (0.30 pts)")
            total_score += 0.30
        elif found_count > 0:
            pts = round(found_count * 0.075, 3)
            print(f"PARTIAL: Component 3 — {found_count}/4 section headings found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 3 — No matching Heading 2 paragraphs found. "
                  f"Found headings: {heading2_texts}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Horizontal lines between sections (0.20 points)
    # Horizontal lines are implemented as paragraphs with bottom border
    # There should be 3 horizontal lines (between 4 sections)
    try:
        hr_count = 0
        for para in doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    bottom = pBdr.find(qn('w:bottom'))
                    if bottom is not None:
                        hr_count += 1

        if hr_count >= 3:
            print(f"PASS: Component 4 — {hr_count} horizontal lines found (0.20 pts)")
            total_score += 0.20
        elif hr_count > 0:
            pts = round(hr_count * (0.20 / 3), 3)
            print(f"PARTIAL: Component 4 — {hr_count}/3 horizontal lines found ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 4 — No horizontal lines (bottom borders) found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Centered page numbers in footer (0.15 points)
    # Check for PAGE field code in footer with center alignment
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        section = doc.sections[0]
        footer = section.footer
        page_field_count = 0
        centered_para_count = 0

        if footer and footer.paragraphs:
            for p in footer.paragraphs:
                # Check alignment
                if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    centered_para_count += 1

                # Check for PAGE field code in runs
                for r in p.runs:
                    for elem in r._element:
                        if elem.tag.endswith('instrText') and 'PAGE' in (elem.text or ''):
                            page_field_count += 1

        if page_field_count > 0 and centered_para_count > 0:
            print(f"PASS: Component 5 — Centered page number field found in footer (0.15 pts)")
            total_score += 0.15
        elif page_field_count > 0 and centered_para_count == 0:
            print(f"PARTIAL: Component 5 — Page number found but not centered (0.07 pts)")
            total_score += 0.07
        else:
            print(f"FAIL: Component 5 — page_fields={page_field_count}, centered_paras={centered_para_count}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
