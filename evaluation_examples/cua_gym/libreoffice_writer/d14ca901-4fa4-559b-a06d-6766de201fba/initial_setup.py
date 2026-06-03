"""
Initial Setup: Disable AutoCorrect (c) -> © replacement
Task ID: writer_edit_068
Domain: libreoffice_writer

Creates programming_notes.docx at ~/Desktop/ with code examples containing (c) references.
The default LibreOffice AutoCorrect is left intact so the agent must disable the (c)->© entry.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_edit_068'
DESKTOP = f'{WORKDIR}/Desktop'
OUTPUT = f'{DESKTOP}/programming_notes.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(DESKTOP, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('C Programming Reference Notes', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    doc.add_paragraph(
        'These notes cover key concepts in C programming with code examples. '
        'Variables and pointers are fundamental to understanding memory management.'
    )

    doc.add_heading('1. Pointer Basics', level=1)
    doc.add_paragraph(
        'In C, a pointer variable stores a memory address. The declaration syntax uses '
        'the asterisk (*) operator. Consider the following example:'
    )

    # Code block style paragraph
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(
        'int *p;      /* p is a pointer to int */\n'
        'int c = 42;  /* c is an integer variable */\n'
        'p = &c;      /* p now holds the address of c */'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)

    doc.add_paragraph(
        'Here, (c) refers to the variable c, not the copyright symbol. '
        'When we write p = &(c), we take the address of the variable (c).'
    )

    doc.add_heading('2. Function Parameters', level=1)
    doc.add_paragraph(
        'Functions can accept pointer parameters to modify the caller\'s data. '
        'The convention (c) is used throughout this codebase to denote a char variable:'
    )

    code_para2 = doc.add_paragraph()
    code_run2 = code_para2.add_run(
        'void process(char *buf, int len) {\n'
        '    char (c) = buf[0];  /* (c) is a local char copy */\n'
        '    if ((c) == \'\\0\') return;\n'
        '    /* process (c) and remaining chars */\n'
        '}'
    )
    code_run2.font.name = 'Courier New'
    code_run2.font.size = Pt(10)

    doc.add_heading('3. Bitwise Operations', level=1)
    doc.add_paragraph(
        'Bitwise operations work on individual bits of integer types. '
        'The expression (c) & 0xFF masks the lower 8 bits of variable (c):'
    )

    code_para3 = doc.add_paragraph()
    code_run3 = code_para3.add_run(
        'unsigned char (c) = input & 0xFF;\n'
        'int result = (c) >> 2;  /* right shift (c) by 2 bits */\n'
        'int masked = (c) & 0x0F;  /* lower nibble of (c) */'
    )
    code_run3.font.name = 'Courier New'
    code_run3.font.size = Pt(10)

    doc.add_heading('4. String Processing', level=1)
    doc.add_paragraph(
        'String processing commonly iterates over each character. '
        'In legacy code, the variable (c) often holds the current character:'
    )

    code_para4 = doc.add_paragraph()
    code_run4 = code_para4.add_run(
        'int count_vowels(const char *str) {\n'
        '    int count = 0;\n'
        '    char (c);\n'
        '    while ((c) = *str++) {\n'
        '        if ((c)==\'a\'||(c)==\'e\'||(c)==\'i\'||(c)==\'o\'||(c)==\'u\')\n'
        '            count++;\n'
        '    }\n'
        '    return count;\n'
        '}'
    )
    code_run4.font.name = 'Courier New'
    code_run4.font.size = Pt(10)

    doc.add_heading('5. Known Issue', level=1)
    issue_para = doc.add_paragraph()
    issue_run = issue_para.add_run('PROBLEM: ')
    issue_run.bold = True
    issue_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    issue_para.add_run(
        'LibreOffice AutoCorrect changes every occurrence of (c) to the copyright '
        'symbol \u00a9 while typing. This corrupts all code examples in this document. '
        'The AutoCorrect entry (c) \u2192 \u00a9 must be removed from '
        'Tools > AutoCorrect Options > Replace tab.'
    )

    doc.add_paragraph(
        'Until this AutoCorrect entry is disabled, the variable references (c) in '
        'new code examples will be incorrectly converted to \u00a9 on input.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
