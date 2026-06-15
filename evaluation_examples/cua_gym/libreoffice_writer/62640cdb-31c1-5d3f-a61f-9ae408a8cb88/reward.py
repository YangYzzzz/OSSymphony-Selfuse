"""
Reward Script: Professional letterhead template with header, body sections, and footer
Task ID: writer_rd_083
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25) - Header: company name bold 18pt green, address/phone 9pt
  Component 2 (0.20) - Header horizontal line (green bottom border)
  Component 3 (0.25) - Body structure: date right-aligned, address block double-spaced,
                        Subject bold, body 1.15 spacing, signature block
  Component 4 (0.15) - Footer: centered website text in gray
  Component 5 (0.15) - Footer horizontal line (gray border)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_083'

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_bottom_border(para):
    """Return bottom border element from paragraph, or None."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        return None
    return pBdr.find(qn('w:bottom'))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    header = section.header
    footer = section.footer

    # =====================================================
    # Component 1: Header content - company name bold 18pt green + address/phone 9pt (0.25 pts)
    # =====================================================
    try:
        header_paras = header.paragraphs
        header_texts = [p.text.strip() for p in header_paras]

        comp1_score = 0.0

        # Check company name "GreenTech Solutions" in header, bold, ~18pt, green color
        found_company = False
        for p in header_paras:
            if 'GreenTech Solutions' in p.text:
                for r in p.runs:
                    if 'GreenTech Solutions' in r.text:
                        is_bold = r.font.bold is True
                        size_ok = r.font.size is not None and abs(r.font.size.pt - 18.0) < 1.0
                        color_ok = r.font.color.rgb is not None and str(r.font.color.rgb) == '006633'
                        if is_bold and size_ok and color_ok:
                            comp1_score += 0.15
                            print(f"PASS: Company name 'GreenTech Solutions' bold 18pt green (0.15 pts)")
                            found_company = True
                        elif is_bold or size_ok:
                            comp1_score += 0.05
                            print(f"PARTIAL: Company name found but formatting incomplete (bold={is_bold}, size_ok={size_ok}, color_ok={color_ok})")
                            found_company = True
                        else:
                            print(f"FAIL: Company name formatting wrong (bold={is_bold}, size={r.font.size}, color={r.font.color.rgb})")
                            found_company = True
                        break
                if found_company:
                    break

        if not found_company:
            print("FAIL: 'GreenTech Solutions' not found in header")

        # Check address and phone in header at ~9pt
        has_address = any('789 Elm Street' in p.text or 'Portland' in p.text for p in header_paras)
        has_phone = any('(503) 555-0199' in p.text or '503' in p.text for p in header_paras)

        addr_size_ok = False
        phone_size_ok = False
        for p in header_paras:
            for r in p.runs:
                if '789 Elm' in r.text or 'Portland' in r.text:
                    if r.font.size is not None and abs(r.font.size.pt - 9.0) < 1.5:
                        addr_size_ok = True
                if '503' in r.text or '555-0199' in r.text:
                    if r.font.size is not None and abs(r.font.size.pt - 9.0) < 1.5:
                        phone_size_ok = True

        if has_address and has_phone:
            comp1_score += 0.10
            print(f"PASS: Address and phone found in header (0.10 pts)")
        elif has_address or has_phone:
            comp1_score += 0.05
            print(f"PARTIAL: Only {'address' if has_address else 'phone'} found in header")
        else:
            print("FAIL: No address or phone in header")

        total_score += comp1_score
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =====================================================
    # Component 2: Header green horizontal line (bottom border on a header paragraph) (0.20 pts)
    # =====================================================
    try:
        found_hdr_border = False
        for p in header.paragraphs:
            border = get_bottom_border(p)
            if border is not None:
                val = border.get(qn('w:val'))
                color = border.get(qn('w:color'))
                sz = border.get(qn('w:sz'))
                if val and val != 'none':
                    found_hdr_border = True
                    # Check if it's green-ish
                    is_green = color is not None and color.lower() in ('006633', '008000', '00ff00', '006600', '003300', '009900')
                    # sz=16 is 2pt (sz units are 1/8 pt)
                    is_thick = sz is not None and int(sz) >= 8  # at least 1pt
                    if is_green and is_thick:
                        total_score += 0.20
                        print(f"PASS: Header bottom border — green ({color}), sz={sz} (0.20 pts)")
                    elif is_green or is_thick:
                        total_score += 0.10
                        print(f"PARTIAL: Header border found but color={color}, sz={sz}")
                    else:
                        total_score += 0.05
                        print(f"PARTIAL: Header border exists but color={color}, sz={sz}")
                    break

        if not found_hdr_border:
            print("FAIL: No horizontal line (bottom border) found in header paragraphs")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =====================================================
    # Component 3: Body structure (0.25 pts)
    #   - Right-aligned date field (0.05)
    #   - Double-spaced address block (0.05)
    #   - Bold "Subject:" line (0.05)
    #   - Body text with ~1.15 line spacing (0.05)
    #   - Signature block with "Sincerely," (0.05)
    # =====================================================
    try:
        body_paras = doc.paragraphs
        comp3_score = 0.0

        # Sub-check 3a: Right-aligned date field
        found_date_right = False
        for p in body_paras:
            text = p.text.strip().lower()
            if 'date' in text or text.startswith('[date'):
                if p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    found_date_right = True
                    break
        if found_date_right:
            comp3_score += 0.05
            print("PASS: Right-aligned date field found (0.05 pts)")
        else:
            print("FAIL: No right-aligned date field found")

        # Sub-check 3b: Double-spaced address block (line_spacing ~2.0)
        double_spaced_count = 0
        for p in body_paras:
            ls = p.paragraph_format.line_spacing
            if ls is not None and abs(float(ls) - 2.0) < 0.1:
                text = p.text.strip().lower()
                if any(kw in text for kw in ['recipient', 'name', 'title', 'company', 'address', 'city', 'state', 'zip', '[']):
                    double_spaced_count += 1
        if double_spaced_count >= 2:
            comp3_score += 0.05
            print(f"PASS: Double-spaced address block found ({double_spaced_count} lines) (0.05 pts)")
        else:
            print(f"FAIL: Expected double-spaced address block, found {double_spaced_count} double-spaced address lines")

        # Sub-check 3c: Bold "Subject:" line
        found_bold_subject = False
        for p in body_paras:
            if 'subject' in p.text.lower():
                for r in p.runs:
                    if 'subject' in r.text.lower() and r.font.bold is True:
                        found_bold_subject = True
                        break
                if found_bold_subject:
                    break
        if found_bold_subject:
            comp3_score += 0.05
            print("PASS: Bold 'Subject:' line found (0.05 pts)")
        else:
            print("FAIL: No bold 'Subject:' line found")

        # Sub-check 3d: Body text with ~1.15 line spacing (must have actual content)
        found_body_spacing = False
        for p in body_paras:
            ls = p.paragraph_format.line_spacing
            if ls is not None and abs(float(ls) - 1.15) < 0.1 and len(p.text.strip()) > 10:
                found_body_spacing = True
                break
        if found_body_spacing:
            comp3_score += 0.05
            print("PASS: Body text with 1.15 line spacing found (0.05 pts)")
        else:
            print("FAIL: No body text with 1.15 line spacing found")

        # Sub-check 3e: Signature block with "Sincerely,"
        found_sincerely = False
        for p in body_paras:
            if 'sincerely' in p.text.lower():
                found_sincerely = True
                break
        if found_sincerely:
            comp3_score += 0.05
            print("PASS: 'Sincerely,' found in signature block (0.05 pts)")
        else:
            print("FAIL: No 'Sincerely,' found in body")

        total_score += comp3_score
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =====================================================
    # Component 4: Footer with centered website text in gray (0.15 pts)
    # =====================================================
    try:
        footer_paras = footer.paragraphs
        found_website = False
        for p in footer_paras:
            text = p.text.strip().lower()
            if 'www.' in text or '.com' in text or 'greentech' in text:
                is_centered = p.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
                is_gray = False
                for r in p.runs:
                    if r.font.color.rgb is not None:
                        c = str(r.font.color.rgb).lower()
                        # Accept various gray tones
                        if c in ('808080', '999999', 'a0a0a0', 'b0b0b0', 'c0c0c0', '666666', '777777', '888888', '909090'):
                            is_gray = True
                            break
                        # General gray check: R==G==B or close
                        r_val = int(c[0:2], 16)
                        g_val = int(c[2:4], 16)
                        b_val = int(c[4:6], 16)
                        if abs(r_val - g_val) < 20 and abs(g_val - b_val) < 20 and r_val > 50:
                            is_gray = True
                            break

                if is_centered and is_gray:
                    total_score += 0.15
                    print(f"PASS: Footer has centered gray website text '{p.text.strip()}' (0.15 pts)")
                    found_website = True
                elif is_centered or is_gray:
                    total_score += 0.07
                    print(f"PARTIAL: Footer website found but centered={is_centered}, gray={is_gray}")
                    found_website = True
                else:
                    total_score += 0.03
                    print(f"PARTIAL: Footer website text found but not centered or gray")
                    found_website = True
                break

        if not found_website:
            print("FAIL: No website text found in footer")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =====================================================
    # Component 5: Footer horizontal line (gray border) (0.15 pts)
    # =====================================================
    try:
        found_ftr_border = False
        for p in footer.paragraphs:
            # Check both top and bottom borders
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                continue
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                continue
            for border_type in ('w:top', 'w:bottom'):
                border = pBdr.find(qn(border_type))
                if border is not None:
                    val = border.get(qn('w:val'))
                    if val and val != 'none':
                        color = border.get(qn('w:color'))
                        found_ftr_border = True
                        # Check gray color
                        is_gray = False
                        if color:
                            c = color.lower()
                            r_val = int(c[0:2], 16)
                            g_val = int(c[2:4], 16)
                            b_val = int(c[4:6], 16)
                            if abs(r_val - g_val) < 20 and abs(g_val - b_val) < 20 and r_val > 50:
                                is_gray = True
                        if is_gray:
                            total_score += 0.15
                            print(f"PASS: Footer border — gray ({color}) (0.15 pts)")
                        else:
                            total_score += 0.07
                            print(f"PARTIAL: Footer border exists but color={color}")
                        break
            if found_ftr_border:
                break

        if not found_ftr_border:
            print("FAIL: No horizontal line (border) found in footer paragraphs")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist app state before verification (LibreOffice may have unsaved changes)
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
persist_app_state()

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
