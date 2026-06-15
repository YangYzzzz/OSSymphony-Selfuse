"""
Reward Script: Header with title and page number
Task ID: writer_tech_031
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Header contains 'CloudSync Admin Guide' text
  Component 2 (0.3): Header has a RIGHT-aligned tab stop
  Component 3 (0.3): Header contains a PAGE field code after a tab (page number on right)
"""

import os

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_TAB_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_031'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    Task: Set up the header to display 'CloudSync Admin Guide' on the left
    and the current page number on the right, separated by a tab.

    We verify:
      1. Header text contains 'CloudSync Admin Guide' (0.4 pts)
      2. Header paragraph has a RIGHT tab stop (0.3 pts)
      3. Header contains a PAGE field code positioned after a tab character (0.3 pts)
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get the first section header
    try:
        section = doc.sections[0]
        header = section.header
        if not header.paragraphs:
            print("FAIL: Header has no paragraphs")
            print("REWARD: 0.0")
            return 0.0
    except Exception as e:
        print(f"CRITICAL: Cannot access header: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the header paragraph that contains actual content
    # (there may be multiple paragraphs; pick the one with text)
    header_para = None
    for para in header.paragraphs:
        if para.text.strip():
            header_para = para
            break
    if header_para is None:
        # No text in any header paragraph
        print("FAIL: No text content found in header paragraphs")
        print("REWARD: 0.0")
        return 0.0

    header_text = header_para.text
    header_xml = header_para._element.xml
    print(f"DEBUG: Header text = {repr(header_text)}")

    # Component 1: Header contains 'CloudSync Admin Guide' (0.4 points)
    try:
        if 'CloudSync Admin Guide' in header_text:
            print(f"PASS: Component 1 -- Header contains 'CloudSync Admin Guide' (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 -- Expected 'CloudSync Admin Guide' in header, found: {repr(header_text)}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Header has a RIGHT tab stop (0.3 points)
    try:
        right_tabs = [ts for ts in header_para.paragraph_format.tab_stops
                      if ts.alignment != WD_TAB_ALIGNMENT.CLEAR
                      and ts.alignment == WD_TAB_ALIGNMENT.RIGHT]
        if len(right_tabs) > 0:
            print(f"  Found RIGHT tab stop at position {right_tabs[0].position}")
            print(f"PASS: Component 2 -- Header has RIGHT tab stop (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 2 -- No RIGHT tab stop found in header paragraph")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Header contains a PAGE field code after a tab (0.3 points)
    # We check the XML: there should be w:instrText containing ' PAGE '
    # AND there should be a w:tab element before the field code section
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        # Check for PAGE field code using any() -- derived from XML element search
        instr_elems = header_para._element.findall('.//w:instrText', ns)
        page_field_found = any(
            instr.text and 'PAGE' in instr.text.upper()
            for instr in instr_elems
        )

        # Check for tab character in the paragraph (w:tab inside w:r, not w:tabs)
        tab_elems = header_para._element.findall('.//w:tab', ns)
        tab_char_found = any(
            el.getparent() is not None and el.getparent().tag.endswith('}r')
            for el in tab_elems
        )

        if page_field_found and tab_char_found:
            print(f"PASS: Component 3 -- Header has PAGE field code after tab separator (0.3 pts)")
            total_score += 0.3
        elif page_field_found and not tab_char_found:
            print(f"PARTIAL: Component 3 -- PAGE field found but no tab separator (0.15 pts)")
            total_score += 0.15
        elif not page_field_found and tab_char_found:
            print(f"FAIL: Component 3 -- Tab found but no PAGE field code")
        else:
            print(f"FAIL: Component 3 -- No PAGE field code and no tab character in header")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
persist_app_state("libreoffice_writer")

# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
