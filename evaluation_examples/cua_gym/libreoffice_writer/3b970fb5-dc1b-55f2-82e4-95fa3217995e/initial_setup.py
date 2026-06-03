"""
Initial Setup: CloudSync Admin Guide with empty header
Task ID: writer_tech_031
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_031'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Enable header but leave it EMPTY
    header = section.header
    header.is_linked_to_previous = False
    # Clear any default content - ensure header paragraph exists but is empty
    for para in header.paragraphs:
        para.text = ""

    # === Document Body: CloudSync Admin Guide ===

    # Title
    title = doc.add_heading("CloudSync Admin Guide", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Enterprise Cloud Synchronization Platform")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    version_para = doc.add_paragraph()
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = version_para.add_run("Version 4.2.1 | Last Updated: March 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")  # spacer

    # Chapter 1: Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "CloudSync is a comprehensive enterprise-grade cloud synchronization platform "
        "designed for organizations that require seamless data replication across multiple "
        "cloud providers and on-premises infrastructure. This administration guide covers "
        "installation, configuration, monitoring, and troubleshooting procedures."
    )
    doc.add_paragraph(
        "CloudSync supports bidirectional synchronization with automatic conflict resolution, "
        "end-to-end encryption, and compliance with SOC 2 Type II, HIPAA, and GDPR requirements. "
        "The platform integrates with AWS S3, Azure Blob Storage, Google Cloud Storage, and "
        "MinIO-compatible object stores."
    )

    # Chapter 2: System Requirements
    doc.add_heading("2. System Requirements", level=1)
    doc.add_heading("2.1 Hardware Requirements", level=2)
    doc.add_paragraph(
        "The following hardware specifications are recommended for production deployments:"
    )

    # Requirements table
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    headers = ["Component", "Minimum", "Recommended"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    req_data = [
        ["CPU", "4 cores (x86_64)", "8+ cores (x86_64)"],
        ["RAM", "8 GB", "32 GB"],
        ["Storage", "100 GB SSD", "500 GB NVMe SSD"],
        ["Network", "1 Gbps", "10 Gbps"],
    ]
    for r, row_data in enumerate(req_data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph("")

    doc.add_heading("2.2 Software Requirements", level=2)
    doc.add_paragraph("CloudSync requires the following software dependencies:", style="List Bullet")
    doc.add_paragraph("Ubuntu 22.04 LTS or RHEL 8.x / 9.x", style="List Bullet")
    doc.add_paragraph("Docker Engine 24.0+ and Docker Compose v2.20+", style="List Bullet")
    doc.add_paragraph("PostgreSQL 15+ (for metadata storage)", style="List Bullet")
    doc.add_paragraph("Redis 7.0+ (for caching and job queues)", style="List Bullet")
    doc.add_paragraph("TLS certificates (Let's Encrypt or internal CA)", style="List Bullet")

    # Chapter 3: Installation
    doc.add_heading("3. Installation", level=1)
    doc.add_paragraph(
        "CloudSync can be deployed using Docker containers or installed directly on bare metal. "
        "The containerized deployment is recommended for most environments as it simplifies "
        "dependency management and upgrades."
    )
    doc.add_heading("3.1 Docker Deployment", level=2)
    doc.add_paragraph(
        "Pull the official CloudSync images from the container registry and configure "
        "the environment variables before starting the services. The deployment uses a "
        "docker-compose.yml file that orchestrates all required services."
    )
    doc.add_paragraph(
        "After deployment, verify the installation by accessing the web dashboard at "
        "https://your-server:8443 and logging in with the default administrator credentials. "
        "Change the default password immediately after first login."
    )

    # Chapter 4: Configuration
    doc.add_heading("4. Configuration", level=1)
    doc.add_heading("4.1 Storage Backends", level=2)
    doc.add_paragraph(
        "Each storage backend requires specific credentials and endpoint configuration. "
        "CloudSync validates connectivity during configuration and will report errors "
        "if the provided credentials lack sufficient permissions."
    )

    # Config table
    config_table = doc.add_table(rows=5, cols=4)
    config_table.style = "Table Grid"
    config_headers = ["Provider", "Auth Method", "Required Permissions", "Default Port"]
    for i, h in enumerate(config_headers):
        cell = config_table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    config_data = [
        ["AWS S3", "IAM Role / Access Key", "s3:GetObject, s3:PutObject, s3:ListBucket", "443"],
        ["Azure Blob", "Service Principal", "Storage Blob Data Contributor", "443"],
        ["Google Cloud", "Service Account JSON", "storage.objects.list, storage.objects.create", "443"],
        ["MinIO", "Access Key / Secret Key", "readwrite policy", "9000"],
    ]
    for r, row_data in enumerate(config_data, 1):
        for c, val in enumerate(row_data):
            config_table.cell(r, c).text = val

    doc.add_paragraph("")

    doc.add_heading("4.2 Sync Policies", level=2)
    doc.add_paragraph(
        "Synchronization policies define how conflicts are resolved and which files are "
        "included or excluded from replication. Policies can be applied globally or per "
        "sync pair. Common conflict resolution strategies include last-writer-wins, "
        "manual resolution, and merge-based reconciliation."
    )

    # Chapter 5: Monitoring and Alerting
    doc.add_heading("5. Monitoring and Alerting", level=1)
    doc.add_paragraph(
        "CloudSync exposes Prometheus-compatible metrics on port 9090. Key metrics to monitor "
        "include sync latency, throughput, error rates, and storage utilization. Grafana dashboards "
        "are included in the deployment bundle for quick visualization setup."
    )
    doc.add_paragraph(
        "Configure alerting rules for critical thresholds such as sync failures exceeding 5% of "
        "total operations, replication lag exceeding 30 minutes, or storage capacity reaching 85% "
        "utilization. Integration with PagerDuty, Slack, and email is supported out of the box."
    )

    # Chapter 6: Troubleshooting
    doc.add_heading("6. Troubleshooting", level=1)
    doc.add_paragraph(
        "When diagnosing sync issues, start by checking the CloudSync logs located at "
        "/var/log/cloudsync/. Common issues include network timeouts, credential expiration, "
        "and storage quota exhaustion. The diagnostic command 'cloudsync doctor' performs "
        "automated health checks across all configured backends."
    )
    doc.add_paragraph(
        "For escalated support, collect the diagnostic bundle using 'cloudsync support-bundle' "
        "and submit it through the enterprise support portal at https://support.cloudsync.io."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
