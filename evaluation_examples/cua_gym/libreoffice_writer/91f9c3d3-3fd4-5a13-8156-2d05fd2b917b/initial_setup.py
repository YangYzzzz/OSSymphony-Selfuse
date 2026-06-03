"""
Initial Setup: Create editing checklist and landscape photo for GIMP editing task
Task ID: osworld_multi_apps_writer_to_gimp_007
Domain: libreoffice_writer + gimp (multi-app)
"""

import os
import shlex
import subprocess
import time
import struct
import zlib

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_to_gimp_007'
CHECKLIST_FILE = f'{WORKDIR}/editing_checklist.docx'
LANDSCAPE_FILE = f'{WORKDIR}/landscape.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_checklist():
    """Create editing_checklist.docx with ordered list of photo editing operations."""
    doc = Document()

    # Title
    title = doc.add_heading("Photo Editing Checklist", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        "Please apply the following editing operations to the landscape photo in the specified order:"
    )
    intro.paragraph_format.space_after = Pt(6)

    # Ordered list of editing steps
    steps = [
        "Crop the image to a 16:9 aspect ratio, centered on the image.",
        "Increase the saturation by 20% (use a saturation enhancement factor of 1.2).",
        "Apply a vignette effect: darken the edges of the image with a gradual circular fade to black.",
    ]

    for step in steps:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.space_after = Pt(4)

    # Footer note
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "Note: Save the final result as 'landscape_edited.jpg' on the Desktop."
    )
    run = note.runs[0]
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(CHECKLIST_FILE)
    print(f'Checklist created: {CHECKLIST_FILE}')


def create_landscape_jpg():
    """Create a realistic landscape photo using Pillow with numpy."""
    try:
        import numpy as np
        from PIL import Image, ImageDraw
    except ImportError:
        print("Pillow/numpy not available, creating minimal JPEG")
        return

    # Image dimensions: 1920x1200 (not 16:9 yet — task requires cropping)
    width, height = 1920, 1200
    img_array = np.zeros((height, width, 3), dtype=np.uint8)

    # --- Sky gradient (top ~40%) ---
    sky_height = int(height * 0.40)
    for y in range(sky_height):
        t = y / sky_height  # 0 at top, 1 at horizon
        # Gradient from deep blue at top to light blue at horizon
        r = int(30 + t * 135)   # 30 -> 165
        g = int(80 + t * 120)   # 80 -> 200
        b = int(180 + t * 55)   # 180 -> 235
        img_array[y, :] = [r, g, b]

    # --- Mountains (middle band) ---
    mountain_base_y = sky_height
    mountain_peak_y = int(height * 0.25)

    # Define multiple mountain peaks
    peaks = [
        (200, mountain_peak_y + 60),
        (420, mountain_peak_y + 20),
        (650, mountain_peak_y + 80),
        (900, mountain_peak_y - 10),
        (1100, mountain_peak_y + 40),
        (1350, mountain_peak_y + 15),
        (1550, mountain_peak_y + 70),
        (1750, mountain_peak_y + 30),
        (1920, mountain_peak_y + 90),
    ]

    img_pil = Image.fromarray(img_array)
    draw = ImageDraw.Draw(img_pil)

    # Draw mountains as polygons
    mountain_color_far = (110, 120, 145)   # far mountains: bluish gray
    mountain_color_near = (80, 95, 75)     # near mountains: dark green-gray

    # Far mountains (lighter)
    for i in range(len(peaks) - 1):
        x1, y1 = peaks[i]
        x2, y2 = peaks[i + 1]
        poly = [
            (x1, mountain_base_y + 80),
            (x1, y1),
            ((x1 + x2) // 2, min(y1, y2) - 20),
            (x2, y2),
            (x2, mountain_base_y + 80),
        ]
        draw.polygon(poly, fill=mountain_color_far)

    # Near mountains (darker, lower)
    near_peaks = [
        (0, int(height * 0.42)),
        (300, int(height * 0.38)),
        (600, int(height * 0.44)),
        (900, int(height * 0.36)),
        (1200, int(height * 0.41)),
        (1500, int(height * 0.37)),
        (1920, int(height * 0.43)),
    ]
    for i in range(len(near_peaks) - 1):
        x1, y1 = near_peaks[i]
        x2, y2 = near_peaks[i + 1]
        poly = [
            (x1, mountain_base_y + 200),
            (x1, y1),
            ((x1 + x2) // 2, min(y1, y2) - 10),
            (x2, y2),
            (x2, mountain_base_y + 200),
        ]
        draw.polygon(poly, fill=mountain_color_near)

    img_array = np.array(img_pil)

    # --- Forest / meadow (lower ~55%) ---
    forest_start = int(height * 0.42)
    for y in range(forest_start, height):
        t = (y - forest_start) / (height - forest_start)  # 0=top of forest, 1=bottom
        # Gradient from dark forest green at top to golden meadow at bottom
        r = int(40 + t * 120)   # 40 -> 160
        g = int(90 + t * 80)    # 90 -> 170
        b = int(35 + t * 30)    # 35 -> 65
        img_array[y, :] = [r, g, b]

    img_pil = Image.fromarray(img_array)
    draw = ImageDraw.Draw(img_pil)

    # Add some tree silhouettes along the treeline
    treeline_y = forest_start + 10
    tree_color = (25, 60, 30)
    for tx in range(50, width - 50, 80):
        import random
        random.seed(tx)
        tree_h = random.randint(40, 90)
        tree_w = random.randint(20, 40)
        # Simple triangular tree
        draw.polygon([
            (tx, treeline_y),
            (tx - tree_w // 2, treeline_y + tree_h),
            (tx + tree_w // 2, treeline_y + tree_h),
        ], fill=tree_color)

    # Add a river/path element
    river_color = (100, 150, 190)
    river_pts = [(width // 2 - 30, treeline_y + 50)]
    x = width // 2 - 30
    for y in range(treeline_y + 50, height, 20):
        x += (y % 3 - 1) * 5
        river_pts.append((x, y))
    if len(river_pts) > 1:
        draw.line(river_pts, fill=river_color, width=8)

    # Add clouds in sky
    cloud_color = (240, 245, 255)
    for cx, cy in [(300, 80), (700, 50), (1100, 100), (1500, 70), (1750, 120)]:
        draw.ellipse([cx - 60, cy - 25, cx + 60, cy + 25], fill=cloud_color)
        draw.ellipse([cx - 40, cy - 40, cx + 40, cy + 10], fill=cloud_color)
        draw.ellipse([cx + 20, cy - 30, cx + 90, cy + 15], fill=cloud_color)

    # Save as JPEG
    img_pil.save(LANDSCAPE_FILE, 'JPEG', quality=92)
    print(f'Landscape photo created: {LANDSCAPE_FILE}  size={img_pil.size}')


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    # Create the checklist document
    create_checklist()

    # Create the landscape photo
    create_landscape_jpg()

    # Verify outputs
    if os.path.exists(CHECKLIST_FILE):
        print(f'  OK: {CHECKLIST_FILE}')
    else:
        print(f'  ERROR: {CHECKLIST_FILE} not found!')

    if os.path.exists(LANDSCAPE_FILE):
        print(f'  OK: {LANDSCAPE_FILE}')
    else:
        print(f'  ERROR: {LANDSCAPE_FILE} not found!')

    # GUI-ready startup: open checklist in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{CHECKLIST_FILE}"', delay_sec=2.0)
    print('GUI_READY: Launched LibreOffice Writer with editing_checklist.docx')


create_initial()
