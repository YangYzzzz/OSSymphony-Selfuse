"""
Initial Setup: Create a thesis-in-progress Writer document with no custom macros.
Task ID: writer_acad_095
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_095'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title page content ---
    title = doc.add_heading('Machine Learning Approaches to Urban Traffic Flow Optimization', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy in Computer Science')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(36)
    run = author.add_run('by\nElena Vasquez-Morrison')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept.paragraph_format.space_before = Pt(24)
    run = dept.add_run('Department of Computer Science\nStanford University\nMarch 2026')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # --- Page break to Abstract ---
    doc.add_page_break()

    doc.add_heading('Abstract', level=1)
    abstract_text = (
        'This thesis presents a comprehensive framework for optimizing urban traffic flow '
        'using deep reinforcement learning techniques combined with graph neural networks. '
        'We propose TrafficNet, a novel architecture that models intersection dependencies '
        'as a dynamic graph and learns adaptive signal control policies through multi-agent '
        'reinforcement learning. Our experiments on real-world traffic datasets from three '
        'major metropolitan areas demonstrate that TrafficNet reduces average commute times '
        'by 18.3% and decreases carbon emissions from idling vehicles by 12.7% compared to '
        'existing adaptive signal control systems. Additionally, we introduce a transfer '
        'learning methodology that enables policies trained on one city to be effectively '
        'deployed in another with minimal fine-tuning, achieving 89.2% of fully-trained '
        'performance after only 48 hours of adaptation.'
    )
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # --- Page break to Chapter 1 ---
    doc.add_page_break()

    doc.add_heading('Chapter 1: Introduction', level=1)

    intro_paras = [
        (
            'Urban traffic congestion remains one of the most pressing challenges facing '
            'modern cities worldwide. According to the Texas A&M Transportation Institute, '
            'the average American commuter spent 54 extra hours in traffic in 2024, costing '
            'the national economy an estimated $87 billion in lost productivity and wasted '
            'fuel (Schrank et al., 2024). These figures have continued to grow despite '
            'significant investments in road infrastructure expansion.'
        ),
        (
            'Traditional approaches to traffic signal control rely on fixed-time plans or '
            'simple actuated controllers that respond to local detector data. While these '
            'methods provide baseline functionality, they fail to capture the complex '
            'spatiotemporal dependencies that characterize urban traffic networks. The '
            'emergence of connected vehicle technology and widespread sensor deployment has '
            'created opportunities for more sophisticated, data-driven approaches.'
        ),
        (
            'Recent advances in deep reinforcement learning (DRL) have shown promising '
            'results in sequential decision-making tasks with high-dimensional state spaces '
            '(Mnih et al., 2015; Silver et al., 2017). Several researchers have applied '
            'DRL to traffic signal control, treating each intersection as an agent that '
            'learns to optimize local traffic flow (Wei et al., 2018; Zheng et al., 2019). '
            'However, these approaches often treat intersections independently, ignoring '
            'the critical dependencies between adjacent signals that can lead to network-level '
            'inefficiencies.'
        ),
    ]

    for text in intro_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.add_heading('1.1 Problem Statement', level=2)

    problem_text = (
        'The central problem addressed in this thesis is how to develop a scalable, '
        'transferable traffic signal control system that jointly optimizes signal timing '
        'across an entire urban network while accounting for the dynamic, stochastic nature '
        'of traffic demand. Specifically, we seek to answer three research questions:'
    )
    p = doc.add_paragraph(problem_text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    rqs = [
        'RQ1: How can graph neural networks effectively capture the spatiotemporal '
        'dependencies between intersections in a traffic network?',
        'RQ2: What multi-agent reinforcement learning framework enables cooperative '
        'signal control that achieves network-level optimization?',
        'RQ3: To what extent can traffic signal control policies be transferred '
        'between cities with different network topologies and traffic patterns?',
    ]
    for rq in rqs:
        p = doc.add_paragraph(rq, style='List Number')
        p.paragraph_format.line_spacing = 2.0
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.add_heading('1.2 Contributions', level=2)

    contributions_text = (
        'This thesis makes the following key contributions to the field of intelligent '
        'transportation systems:'
    )
    p = doc.add_paragraph(contributions_text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    contribs = [
        'TrafficNet Architecture: A novel deep learning architecture that combines '
        'graph attention networks with multi-agent proximal policy optimization to '
        'learn cooperative traffic signal control policies.',
        'Real-World Evaluation: Comprehensive experiments using calibrated simulation '
        'models of downtown San Francisco, central Tokyo, and inner London, validating '
        'our approach against five baseline methods.',
        'Transfer Learning Framework: A domain adaptation methodology for traffic '
        'control that reduces deployment costs by enabling cross-city policy transfer '
        'with minimal retraining.',
    ]
    for c in contribs:
        p = doc.add_paragraph(c, style='List Bullet')
        p.paragraph_format.line_spacing = 2.0
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # --- Page break to Chapter 2 ---
    doc.add_page_break()

    doc.add_heading('Chapter 2: Literature Review', level=1)

    lit_review_paras = [
        (
            'This chapter provides a comprehensive review of the relevant literature spanning '
            'three key areas: traditional traffic signal control systems, reinforcement learning '
            'for transportation, and graph neural networks for spatiotemporal modeling.'
        ),
        (
            'The history of traffic signal control dates back to the late 1960s with the '
            'development of TRANSYT (Robertson, 1969), which optimized fixed-time signal '
            'plans using a macroscopic traffic model. This was followed by adaptive systems '
            'such as SCOOT (Hunt et al., 1981) and SCATS (Lowrie, 1990), which adjusted '
            'signal parameters in real-time based on detector measurements. While these '
            'systems represented significant advances, their reliance on hand-crafted '
            'optimization rules and simplified traffic models limited their effectiveness '
            'in highly dynamic conditions.'
        ),
        (
            'The application of reinforcement learning to traffic control began with '
            'Thorpe and Anderson (1996), who used a simple Q-learning agent to control '
            'a single intersection. Subsequent work by Abdulhai et al. (2003) extended '
            'this to multi-phase intersections, demonstrating the potential of model-free '
            'learning approaches. The deep reinforcement learning revolution, initiated '
            'by the DQN algorithm (Mnih et al., 2015), enabled researchers to handle '
            'the high-dimensional observation spaces typical of traffic networks.'
        ),
    ]

    for text in lit_review_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # --- Page break to Chapter 3 ---
    doc.add_page_break()

    doc.add_heading('Chapter 3: Methodology', level=1)

    method_text = (
        'This chapter describes the TrafficNet architecture and the multi-agent reinforcement '
        'learning framework used for cooperative traffic signal control. We first formalize '
        'the traffic signal control problem as a decentralized partially observable Markov '
        'decision process (Dec-POMDP), then detail the graph neural network component and '
        'the policy optimization algorithm.'
    )
    p = doc.add_paragraph(method_text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_heading('3.1 Problem Formulation', level=2)

    formulation_text = (
        'We model the traffic network as a directed graph G = (V, E), where V represents '
        'the set of signalized intersections and E represents road segments connecting them. '
        'Each intersection i maintains a local observation comprising queue lengths, waiting '
        'times, and current phase information. The global state is the joint observation of '
        'all intersections, but each agent only has access to observations from its local '
        'neighborhood defined by the graph structure.'
    )
    p = doc.add_paragraph(formulation_text)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # --- References placeholder ---
    doc.add_page_break()
    doc.add_heading('References', level=1)

    refs = [
        'Abdulhai, B., Pringle, R., & Karakoulas, G. J. (2003). Reinforcement learning for true adaptive traffic signal control. Journal of Transportation Engineering, 129(3), 278-285.',
        'Hunt, P. B., Robertson, D. I., Bretherton, R. D., & Royle, M. C. (1981). The SCOOT on-line traffic signal optimisation technique. Traffic Engineering & Control, 23(4), 190-192.',
        'Lowrie, P. R. (1990). SCATS: Sydney co-ordinated adaptive traffic system. ATEC, Sydney.',
        'Mnih, V., Kavukcuoglu, K., Silver, D., et al. (2015). Human-level control through deep reinforcement learning. Nature, 518(7540), 529-533.',
        'Robertson, D. I. (1969). TRANSYT: A traffic network study tool. RRL Report LR 253.',
        'Schrank, D., Eisele, B., & Lomax, T. (2024). 2024 Urban Mobility Report. Texas A&M Transportation Institute.',
        'Silver, D., Schrittwieser, J., Simonyan, K., et al. (2017). Mastering the game of Go without human knowledge. Nature, 550(7676), 354-359.',
        'Thorpe, T. L., & Anderson, C. W. (1996). Traffic light control using SARSA with three state representations. Technical Report, Colorado State University.',
        'Wei, H., Zheng, G., Yao, H., & Li, Z. (2018). IntelliLight: A reinforcement learning approach for intelligent traffic light control. In KDD 2018.',
        'Zheng, G., Xiong, Y., Zang, X., et al. (2019). Learning phase competition for traffic signal control. In CIKM 2019.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Ensure default macro state: reset Module1.xba to empty
    module1_path = '/home/user/.config/libreoffice/4/user/basic/Standard/Module1.xba'
    default_module1 = '''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE script:module PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "module.dtd">
<script:module xmlns:script="http://openoffice.org/2000/script" script:name="Module1" script:language="StarBasic">REM  *****  BASIC  *****

Sub Main

End Sub</script:module>'''
    os.makedirs(os.path.dirname(module1_path), exist_ok=True)
    with open(module1_path, 'w') as f:
        f.write(default_module1)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
