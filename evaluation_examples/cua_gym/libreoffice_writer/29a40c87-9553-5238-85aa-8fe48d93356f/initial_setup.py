"""
Initial Setup: Writer to GIMP multi-app photo processing task
Task ID: osworld_multi_apps_writer_to_gimp_010
Domain: multi_apps (libreoffice_writer + gimp)

Creates:
  - /home/user/Desktop/processing_request.docx  — instructions document
  - /home/user/Desktop/event_photo.jpg          — source photo to process
  Does NOT create event_photo_processed.jpg (that is the agent's task)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
import io

WORKDIR = '/home/user/Desktop'
TASK_ID = 'osworld_multi_apps_writer_to_gimp_010'
DOC_PATH = f'{WORKDIR}/processing_request.docx'
PHOTO_PATH = f'{WORKDIR}/event_photo.jpg'
# Output (what agent must produce) — must NOT exist in initial state
OUTPUT_PATH = f'{WORKDIR}/event_photo_processed.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_processing_request_doc():
    """Create a realistic photo processing instructions document."""
    doc = Document()

    # Title
    title = doc.add_heading('Photo Processing Request', level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        'Please process the provided event photo according to the following specifications. '
        'All steps must be applied in the order listed and saved to the specified output file.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Processing steps section heading
    doc.add_heading('Processing Instructions', level=2)

    # Step 1: Resize
    step1_heading = doc.add_paragraph()
    run = step1_heading.add_run('Step 1: Resize Image')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(
        'Resize the image to exactly 1920 x 1080 pixels. '
        'Do not maintain aspect ratio — stretch or compress as needed to fit the target dimensions.'
    )

    # Step 2: Brightness
    step2_heading = doc.add_paragraph()
    run = step2_heading.add_run('Step 2: Increase Brightness')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(
        'Increase the overall brightness of the image by a factor of 1.3 '
        '(i.e., apply a brightness enhancement factor of 1.3 using the Brightness tool).'
    )

    # Step 3: Watermark
    step3_heading = doc.add_paragraph()
    run = step3_heading.add_run('Step 3: Add Text Watermark')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(
        'Add the watermark text "Event 2025" in the bottom-right corner of the image. '
        'Use a white, semi-transparent text style so it is visible but does not obscure the photo.'
    )

    # Output section
    doc.add_heading('Output Requirements', level=2)
    output_para = doc.add_paragraph()
    run1 = output_para.add_run('Save the processed image as: ')
    run2 = output_para.add_run('event_photo_processed.jpg')
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x20, 0x6A, 0xC9)
    run3 = output_para.add_run(' on the Desktop.')

    # Notes
    doc.add_heading('Notes', level=2)
    doc.add_paragraph(
        'The source file is event_photo.jpg located on the Desktop. '
        'Please ensure all three processing steps are applied before saving.'
    )

    # Additional context paragraph
    notes_list = [
        'Use GIMP Script-Fu console or manual tools for processing.',
        'The output file must be in JPEG format.',
        'Quality setting: 90% (default JPEG quality is acceptable).',
    ]
    for item in notes_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(DOC_PATH)
    print(f'Document created: {DOC_PATH}')


def create_event_photo():
    """Create a realistic-looking event photo as the source image."""
    # Create a 2560x1440 image to simulate a high-res photo taken at an event
    width, height = 2560, 1440
    img = Image.new('RGB', (width, height))

    # Draw a gradient sky background
    pixels = img.load()
    for y in range(height):
        r = int(100 + (y / height) * 80)
        g = int(140 + (y / height) * 50)
        b = int(200 - (y / height) * 60)
        for x in range(width):
            pixels[x, y] = (r, g, b)

    draw = ImageDraw.Draw(img)

    # Ground / stage area
    draw.rectangle([0, int(height * 0.65), width, height], fill=(60, 45, 30))

    # Stage platform
    draw.rectangle([int(width * 0.15), int(height * 0.55), int(width * 0.85), int(height * 0.72)],
                   fill=(80, 65, 50))
    draw.rectangle([int(width * 0.15), int(height * 0.55), int(width * 0.85), int(height * 0.575)],
                   fill=(110, 90, 70))

    # Stage backdrop panels
    panel_colors = [(180, 80, 50), (60, 120, 180), (160, 160, 50), (80, 160, 80), (160, 60, 160)]
    panel_w = int((width * 0.7) / len(panel_colors))
    for i, color in enumerate(panel_colors):
        x0 = int(width * 0.15) + i * panel_w
        draw.rectangle([x0, int(height * 0.2), x0 + panel_w - 4, int(height * 0.55)], fill=color)

    # Crowd silhouettes
    import random
    random.seed(42)
    for i in range(120):
        cx = random.randint(50, width - 50)
        cy = random.randint(int(height * 0.68), int(height * 0.88))
        head_r = random.randint(14, 24)
        body_h = random.randint(40, 70)
        shade = random.randint(20, 50)
        draw.ellipse([cx - head_r, cy - head_r, cx + head_r, cy + head_r],
                     fill=(shade, shade, shade))
        draw.rectangle([cx - head_r + 4, cy + head_r - 2, cx + head_r - 4, cy + head_r + body_h],
                       fill=(shade + 5, shade + 5, shade + 5))

    # Spotlights / stage lights
    light_positions = [(int(width * 0.25), int(height * 0.08)),
                       (int(width * 0.5), int(height * 0.05)),
                       (int(width * 0.75), int(height * 0.08))]
    for lx, ly in light_positions:
        draw.ellipse([lx - 18, ly - 18, lx + 18, ly + 18], fill=(255, 255, 200))
        # Light cone
        cone_pts = [(lx, ly + 18),
                    (lx - 80, int(height * 0.52)),
                    (lx + 80, int(height * 0.52))]
        draw.polygon(cone_pts, fill=(255, 255, 180, 40))

    # Event banner text at top
    try:
        font_large = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 64)
        font_medium = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 40)
    except Exception:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()

    # Banner background
    draw.rectangle([int(width * 0.2), int(height * 0.04), int(width * 0.8), int(height * 0.13)],
                   fill=(20, 20, 60))
    draw.text((int(width * 0.5) - 200, int(height * 0.055)),
              'ANNUAL TECH SUMMIT 2025', fill=(255, 220, 50), font=font_large)

    # Date line
    draw.text((int(width * 0.5) - 160, int(height * 0.11)),
              'March 15, 2025  |  San Francisco Convention Center',
              fill=(200, 200, 200), font=font_medium)

    img.save(PHOTO_PATH, 'JPEG', quality=92)
    print(f'Event photo created: {PHOTO_PATH}  size={img.size}')


def remove_output_if_exists():
    """Ensure the output file does not exist in initial state."""
    if os.path.exists(OUTPUT_PATH):
        os.remove(OUTPUT_PATH)
        print(f'Removed pre-existing output file: {OUTPUT_PATH}')


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    create_processing_request_doc()
    create_event_photo()
    remove_output_if_exists()

    # GUI-ready startup: open the document in Writer and the photo in GIMP
    launch_gui(f'libreoffice --writer "{DOC_PATH}"', delay_sec=3.0)
    launch_gui(f'gimp "{PHOTO_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer and GIMP with DISPLAY=:0')


create_initial()
