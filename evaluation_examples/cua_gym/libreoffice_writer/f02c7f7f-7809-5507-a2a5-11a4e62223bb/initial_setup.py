"""
Initial Setup: Set the second paragraph of this document to double line spacing.
Task ID: osworld_writer_line_spacing_per_paragraph_002
Domain: libreoffice_writer

Creates a 4-paragraph meeting minutes document where ALL paragraphs
use single line spacing (1.0). The task is for the agent to change
only the second paragraph to double line spacing.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Paragraph 1: Meeting title / intro
    p1 = doc.add_paragraph(
        "Meeting Minutes — Project Horizon Quarterly Review"
    )
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_after = Pt(6)

    # Paragraph 2: Attendees (this is the one the agent must change to double spacing)
    p2 = doc.add_paragraph(
        "Attendees: Rachel Nguyen (Project Lead), David Kim (Engineering), "
        "Sandra Okafor (Design), Marcus Ellison (QA), Priya Sharma (Product)."
    )
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_after = Pt(6)

    # Paragraph 3: Discussion summary
    p3 = doc.add_paragraph(
        "Discussion: The team reviewed the Q2 roadmap milestones and identified "
        "three critical blockers: API integration delays with the third-party vendor, "
        "insufficient test coverage on the authentication module, and unresolved "
        "UX feedback from the usability study conducted on March 18, 2025. "
        "David Kim committed to delivering a revised integration plan by April 5, "
        "while Sandra Okafor will present updated wireframes at the next sprint review."
    )
    p3.paragraph_format.line_spacing = 1.0
    p3.paragraph_format.space_after = Pt(6)

    # Paragraph 4: Action items / next steps
    p4 = doc.add_paragraph(
        "Next Steps: A follow-up meeting is scheduled for April 10, 2025 at 10:00 AM. "
        "All action items must be logged in the project tracker before end of business "
        "on April 3, 2025. Marcus Ellison will circulate the updated test plan by "
        "April 7, 2025, and Priya Sharma will prepare the stakeholder briefing deck "
        "for the executive review on April 14, 2025."
    )
    p4.paragraph_format.line_spacing = 1.0
    p4.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
