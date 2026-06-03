"""
Reward Script: Set the second paragraph of this document to double line spacing.
Task ID: osworld_writer_line_spacing_per_paragraph_002
Domain: libreoffice_writer
Scoring:
  Component 1: Second paragraph (index 1) has double line spacing (2.0) — 0.6 pts
  Component 2: Scoped correctly — only para[1] changed; other paragraphs remain single — 0.4 pts
               (Component 2 only awards points if Component 1 passes, ensuring the sub-check
               measures task scope precision, not a pre-existing property.)
"""

import os

from docx import Document
from docx.enum.text import WD_LINE_SPACING

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_002'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have exactly 4 paragraphs
    if len(doc.paragraphs) != 4:
        print(f"PRECONDITION FAIL: Expected 4 paragraphs, found {len(doc.paragraphs)}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Second paragraph (index 1) has double line spacing (2.0) — 0.6 pts
    # This FAILS on initial_env (all paragraphs are single) and PASSES on golden_env (para[1] = double)
    para1_is_double = False
    try:
        para1 = doc.paragraphs[1]
        pf = para1.paragraph_format
        ls = pf.line_spacing
        rule = pf.line_spacing_rule
        # Double line spacing: line_spacing == 2.0 (float) or line_spacing_rule == DOUBLE
        para1_is_double = (ls is not None and abs(float(ls) - 2.0) < 0.05) or (rule == WD_LINE_SPACING.DOUBLE)
        if para1_is_double:
            print(f"PASS: Component 1 — second paragraph has double line spacing (line_spacing={ls}, rule={rule}) (0.6 pts)")
            total_score += 0.6
        else:
            print(f"FAIL: Component 1 — expected double line spacing (2.0), found line_spacing={ls}, rule={rule}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: All other paragraphs (indices 0, 2, 3) still have single line spacing — 0.4 pts
    # This compound check ensures the change is scoped only to para[1].
    # It is ONLY awarded when Component 1 passes (para[1] is double), meaning it verifies that
    # the agent changed ONLY para[1] and not others. Without Component 1 passing, this check
    # would be a pre-existing property of the initial state and MUST NOT earn points.
    if para1_is_double:
        try:
            wrong_paras = []
            details = []
            for idx in [0, 2, 3]:
                para = doc.paragraphs[idx]
                pf = para.paragraph_format
                ls = pf.line_spacing
                rule = pf.line_spacing_rule
                # Single line spacing: line_spacing == 1.0 or None (inherit) or rule == SINGLE
                is_single = (
                    (ls is None) or
                    (abs(float(ls) - 1.0) < 0.05) or
                    (rule == WD_LINE_SPACING.SINGLE)
                )
                if not is_single:
                    wrong_paras.append(f"para[{idx}] line_spacing={ls}, rule={rule}")
                    details.append(f"para[{idx}] line_spacing={ls}, rule={rule}")
                else:
                    details.append(f"para[{idx}] OK (ls={ls})")

            if len(wrong_paras) == 0:
                print(f"PASS: Component 2 — scope correct, non-target paragraphs retain single: {details} (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 2 — non-target paragraphs have wrong spacing: {wrong_paras}")
        except Exception as e:
            print(f"ERROR: Component 2 — {e}")
    else:
        print("SKIP: Component 2 — skipped because Component 1 failed (para[1] not double)")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in the VM
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
