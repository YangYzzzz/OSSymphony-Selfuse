"""
Reward Script: Set only top and bottom borders of table to 2pt solid black
Task ID: writer_tm_032
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): Top border of row 0 cells is 2pt solid black
  Component 2 (0.30): Bottom border of last row cells is 2pt solid black
  Component 3 (0.25): All inner borders and left/right outer borders are removed (none/nil)
  Component 4 (0.15): Cell content unchanged from original
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_032'

# Expected cell content (from task context — 4 cols x 6 rows)
EXPECTED_CONTENT = [
    ['Product Line', 'Q1 Revenue', 'Q2 Revenue', 'Q3 Revenue'],
    ['Enterprise Cloud', '$1,245,300', '$1,389,750', '$1,502,400'],
    ['Developer Tools', '$823,600', '$891,200', '$945,800'],
    ['Data Analytics', '$567,400', '$612,300', '$678,900'],
    ['Security Suite', '$412,800', '$438,500', '$491,200'],
    ['Mobile Platform', '$298,100', '$325,600', '$367,400'],
]


def get_border_info(tc_element, side):
    """
    Get border info for a cell on the given side ('top','bottom','left','right').
    Returns (val, sz) where val is the border style and sz is the size attribute.
    Returns (None, None) if no border element found.
    """
    tcPr = tc_element.find(qn('w:tcPr'))
    if tcPr is None:
        return (None, None)
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        return (None, None)
    side_elem = borders.find(qn(f'w:{side}'))
    if side_elem is None:
        return (None, None)
    val = side_elem.get(qn('w:val'))
    sz = side_elem.get(qn('w:sz'))
    return (val, sz)


def is_border_explicitly_none(tc_element, side):
    """
    Check if a cell border on the given side is EXPLICITLY removed (val='none' or 'nil').
    Absent borders (no element) are NOT considered removed -- they inherit from the table style.
    This distinction is critical: initial file uses TableGrid style with inherited borders,
    while golden file explicitly sets val='none' to remove them.
    """
    val, sz = get_border_info(tc_element, side)
    if val is None:
        return False  # absent = inherited from style, NOT removed
    if val in ('none', 'nil'):
        return True
    if sz is not None and sz == '0' and val != 'single':
        return True
    return False


def is_border_2pt_solid_black(tc_element, side):
    """
    Check if a cell border on the given side is 2pt solid black.
    In OOXML, sz is in 1/8 pt, so 2pt = sz=16.
    """
    val, sz = get_border_info(tc_element, side)
    if val != 'single':
        return False
    if sz is None:
        return False
    # 2pt = 16 eighth-points. Allow some tolerance (14-18 to account for rounding).
    sz_int = int(sz)
    if sz_int < 14 or sz_int > 18:
        return False
    # Check color is black (000000 or auto)
    tcPr = tc_element.find(qn('w:tcPr'))
    borders = tcPr.find(qn('w:tcBorders'))
    side_elem = borders.find(qn(f'w:{side}'))
    color = side_elem.get(qn('w:color'))
    if color is not None and color.lower() not in ('000000', 'auto'):
        return False
    return True


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    if num_rows < 2 or num_cols < 1:
        print(f"FAIL: Table too small ({num_rows}x{num_cols})")
        print("REWARD: 0.0")
        return 0.0

    last_row_idx = num_rows - 1

    # Component 1: Top border of row 0 cells is 2pt solid black (0.30 points)
    try:
        pass_count = 0
        for c in range(num_cols):
            tc = table.cell(0, c)._tc
            if is_border_2pt_solid_black(tc, 'top'):
                pass_count += 1
        ratio = pass_count / num_cols
        if ratio >= 0.75:
            pts = 0.30
            print(f"PASS: Component 1 — Top border of row 0: {pass_count}/{num_cols} cells have 2pt solid black top border (0.30 pts)")
            total_score += pts
        elif ratio > 0:
            pts = round(0.30 * ratio, 2)
            print(f"PARTIAL: Component 1 — Top border of row 0: {pass_count}/{num_cols} cells ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 1 — Top border of row 0: 0/{num_cols} cells have 2pt solid black top border")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Bottom border of last row cells is 2pt solid black (0.30 points)
    try:
        pass_count = 0
        for c in range(num_cols):
            tc = table.cell(last_row_idx, c)._tc
            if is_border_2pt_solid_black(tc, 'bottom'):
                pass_count += 1
        ratio = pass_count / num_cols
        if ratio >= 0.75:
            pts = 0.30
            print(f"PASS: Component 2 — Bottom border of last row: {pass_count}/{num_cols} cells have 2pt solid black bottom border (0.30 pts)")
            total_score += pts
        elif ratio > 0:
            pts = round(0.30 * ratio, 2)
            print(f"PARTIAL: Component 2 — Bottom border of last row: {pass_count}/{num_cols} cells ({pts} pts)")
            total_score += pts
        else:
            print(f"FAIL: Component 2 — Bottom border of last row: 0/{num_cols} cells have 2pt solid black bottom border")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All inner borders and left/right outer borders are removed (0.25 points)
    # Check: for every cell, left and right borders should be none.
    # For row 0 cells: bottom should be none. For last row cells: top should be none.
    # For middle row cells: top and bottom should be none.
    try:
        total_checks = 0
        pass_checks = 0
        for r in range(num_rows):
            for c in range(num_cols):
                tc = table.cell(r, c)._tc
                # Left border: always should be none
                total_checks += 1
                if is_border_explicitly_none(tc, 'left'):
                    pass_checks += 1

                # Right border: always should be none
                total_checks += 1
                if is_border_explicitly_none(tc, 'right'):
                    pass_checks += 1

                # Top border: should be none for all rows except row 0
                if r > 0:
                    total_checks += 1
                    if is_border_explicitly_none(tc, 'top'):
                        pass_checks += 1

                # Bottom border: should be none for all rows except last row
                if r < last_row_idx:
                    total_checks += 1
                    if is_border_explicitly_none(tc, 'bottom'):
                        pass_checks += 1

        if total_checks > 0:
            ratio = pass_checks / total_checks
            if ratio >= 0.90:
                pts = 0.25
                print(f"PASS: Component 3 — Inner/side borders removed: {pass_checks}/{total_checks} checks pass (0.25 pts)")
                total_score += pts
            elif ratio > 0.5:
                pts = round(0.25 * ratio, 2)
                print(f"PARTIAL: Component 3 — Inner/side borders: {pass_checks}/{total_checks} checks pass ({pts} pts)")
                total_score += pts
            else:
                print(f"FAIL: Component 3 — Inner/side borders NOT removed: only {pass_checks}/{total_checks} checks pass")
        else:
            print("FAIL: Component 3 — No border checks possible")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Cell content unchanged (0.15 points)
    try:
        if num_rows == len(EXPECTED_CONTENT) and num_cols == len(EXPECTED_CONTENT[0]):
            match_count = 0
            total_cells = num_rows * num_cols
            for r in range(num_rows):
                for c in range(num_cols):
                    actual = table.cell(r, c).text.strip()
                    expected = EXPECTED_CONTENT[r][c]
                    if actual == expected:
                        match_count += 1
            if match_count == total_cells and total_score > 0:
                print(f"PASS: Component 4 — All {total_cells} cell contents match expected values (0.15 pts)")
                total_score += 0.15
            elif match_count == total_cells and total_score == 0:
                print(f"FAIL: Component 4 — Content intact but no border changes detected (0 pts, gated on border components)")
            else:
                # Partial: only if most cells match and at least one border component scored
                if total_score > 0 and match_count > total_cells * 0.5:
                    pts = round(0.15 * (match_count / total_cells), 2)
                    print(f"PARTIAL: Component 4 — {match_count}/{total_cells} cells match ({pts} pts)")
                    total_score += pts
                else:
                    print(f"FAIL: Component 4 — Content mismatch or no border changes: {match_count}/{total_cells} cells match")
        else:
            print(f"FAIL: Component 4 — Table dimensions {num_rows}x{num_cols} != expected 6x4")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {final_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: attempt to save any unsaved LibreOffice state
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
