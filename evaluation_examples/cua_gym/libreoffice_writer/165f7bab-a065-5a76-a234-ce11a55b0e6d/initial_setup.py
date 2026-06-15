"""
Initial Setup: Set only the top and bottom borders of a minimalist table to 2pt solid black
Task ID: writer_tm_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell using XML. Each border arg is a dict with sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    else:
        # Clear existing
        for child in list(tcBorders):
            tcBorders.remove(child)

    for side_name, border_spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if border_spec is not None:
            el = parse_xml(
                f'<w:{side_name} {nsdecls("w")} w:val="{border_spec["val"]}" '
                f'w:sz="{border_spec["sz"]}" w:space="0" w:color="{border_spec["color"]}"/>'
            )
            tcBorders.append(el)


def create_initial():
    doc = Document()

    # Add a title paragraph
    heading = doc.add_heading('Quarterly Performance Report', level=1)

    intro = doc.add_paragraph(
        'The following table summarizes the quarterly revenue performance '
        'across our four main product lines for the fiscal year 2025.'
    )

    # Create 4-column x 6-row table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'  # This gives all borders by default

    # Populate with realistic data
    headers = ['Product Line', 'Q1 Revenue', 'Q2 Revenue', 'Q3 Revenue']
    data = [
        ['Enterprise Cloud', '$1,245,300', '$1,389,750', '$1,502,400'],
        ['Developer Tools', '$823,600', '$891,200', '$945,800'],
        ['Data Analytics', '$567,400', '$612,300', '$678,900'],
        ['Security Suite', '$412,800', '$438,500', '$491,200'],
        ['Mobile Platform', '$298,100', '$325,600', '$367,400'],
    ]

    # Set headers
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(11)

    # Set data rows
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value

    # Now override all cell borders to 0.5pt solid black (sz=4 in half-points, 0.5pt = 4 eighth-points)
    # In OOXML, sz is in eighth-points, so 0.5pt = 4
    border_half_pt = {"val": "single", "sz": "4", "color": "000000"}

    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(
                cell,
                top=border_half_pt,
                bottom=border_half_pt,
                left=border_half_pt,
                right=border_half_pt,
            )

    # Add a closing paragraph
    doc.add_paragraph(
        'Note: All figures are in USD. Q4 projections will be available in the next report.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
