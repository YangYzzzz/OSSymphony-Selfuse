"""
Initial Setup: Financial approval form - 4 pages, Authorization section on page 4
Task ID: writer_fp_034
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_034'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    return h


def add_body_text(doc, text, bold=False, size=Pt(11)):
    """Add a body paragraph with optional formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = size
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    return para


def add_table_row_data(table, row_idx, values):
    """Populate a table row with values."""
    for col_idx, val in enumerate(values):
        table.cell(row_idx, col_idx).text = str(val)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # === PAGE 1: Cover / Title Page ===
    doc.add_paragraph()  # spacing
    doc.add_paragraph()

    title = doc.add_heading('Capital Expenditure Approval Request', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Meridian Technologies, Inc.')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.bold = True

    doc.add_paragraph()

    info_items = [
        ('Project:', 'Enterprise Data Center Expansion - Phase II'),
        ('Department:', 'Information Technology Infrastructure'),
        ('Requested Amount:', '$2,450,000.00'),
        ('Date Submitted:', 'March 15, 2025'),
        ('Reference No.:', 'CAPEX-2025-0347'),
    ]

    for label, value in info_items:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r1 = para.add_run(label + ' ')
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = para.add_run(value)
        r2.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    conf = doc.add_paragraph()
    conf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = conf.add_run('CONFIDENTIAL - FOR INTERNAL USE ONLY')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r.bold = True

    # === PAGE 2: Executive Summary & Budget ===
    doc.add_page_break()

    add_heading_styled(doc, '1. Executive Summary', level=1)

    add_body_text(doc,
        'This capital expenditure request seeks approval for the second phase of the '
        'Enterprise Data Center Expansion project at our Westfield campus facility. '
        'The expansion is critical to support the projected 40% increase in computing '
        'demand driven by our growing cloud services portfolio and new AI/ML workloads.')

    add_body_text(doc,
        'Phase I, completed in Q3 2024 at a cost of $1.8M (under budget by 7%), '
        'successfully delivered 200 additional rack units and upgraded the primary '
        'cooling infrastructure. Phase II will add 150 rack units, implement redundant '
        'power distribution, and deploy next-generation network switching fabric.')

    add_heading_styled(doc, '2. Budget Breakdown', level=1)

    # Budget table
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'

    headers = ['Category', 'Description', 'Estimated Cost', 'Contingency (10%)']
    add_table_row_data(table, 0, headers)
    # Bold headers
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    budget_data = [
        ['Server Hardware', 'Dell PowerEdge R760 (48 units)', '$864,000', '$86,400'],
        ['Network Equipment', 'Cisco Nexus 9300 switches (8 units)', '$312,000', '$31,200'],
        ['Power Distribution', 'Redundant PDU + UPS upgrade', '$425,000', '$42,500'],
        ['Cooling Systems', 'In-row cooling units (12 units)', '$288,000', '$28,800'],
        ['Cabling & Infrastructure', 'Fiber optic, Cat6a, cable management', '$156,000', '$15,600'],
        ['Installation & Labor', 'Professional services, project mgmt', '$245,000', '$24,500'],
        ['TOTAL', '', '$2,290,000', '$229,000'],
    ]
    for i, row_data in enumerate(budget_data, 1):
        add_table_row_data(table, i, row_data)

    # Bold the TOTAL row
    for cell in table.rows[7].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()
    add_body_text(doc,
        'Total project budget including contingency: $2,519,000. Requested approval '
        'amount of $2,450,000 reflects negotiated vendor discounts not yet reflected '
        'in the line items above.')

    # === PAGE 3: Timeline & Risk Assessment ===
    doc.add_page_break()

    add_heading_styled(doc, '3. Implementation Timeline', level=1)

    timeline_table = doc.add_table(rows=7, cols=4)
    timeline_table.style = 'Table Grid'

    add_table_row_data(timeline_table, 0, ['Phase', 'Activity', 'Start Date', 'End Date'])
    for cell in timeline_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    timeline_data = [
        ['1', 'Procurement & vendor selection', 'Apr 1, 2025', 'May 15, 2025'],
        ['2', 'Site preparation & power upgrade', 'May 1, 2025', 'Jun 30, 2025'],
        ['3', 'Equipment delivery & staging', 'Jun 15, 2025', 'Jul 31, 2025'],
        ['4', 'Installation & rack deployment', 'Jul 1, 2025', 'Sep 15, 2025'],
        ['5', 'Network configuration & testing', 'Sep 1, 2025', 'Oct 15, 2025'],
        ['6', 'Migration & go-live', 'Oct 1, 2025', 'Nov 30, 2025'],
    ]
    for i, row_data in enumerate(timeline_data, 1):
        add_table_row_data(timeline_table, i, row_data)

    doc.add_paragraph()

    add_heading_styled(doc, '4. Risk Assessment', level=1)

    risks = [
        ('Supply Chain Delays', 'Medium',
         'Current lead times for server hardware are 8-12 weeks. Mitigation: Early '
         'procurement initiation and secondary vendor agreements.'),
        ('Power Capacity Constraints', 'High',
         'Existing utility feed may require upgrade for full Phase II load. Mitigation: '
         'Utility company engagement already initiated; temporary generator backup planned.'),
        ('Budget Overrun', 'Low',
         'Phase I came in under budget. 10% contingency built into estimates. '
         'Monthly cost tracking with variance reporting to CFO.'),
        ('Operational Disruption', 'Medium',
         'Active datacenter operations during construction. Mitigation: Weekend and '
         'off-hours work scheduling, isolated power circuits for new equipment.'),
    ]

    risk_table = doc.add_table(rows=len(risks) + 1, cols=3)
    risk_table.style = 'Table Grid'
    add_table_row_data(risk_table, 0, ['Risk Factor', 'Severity', 'Mitigation Strategy'])
    for cell in risk_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i, (risk, severity, mitigation) in enumerate(risks, 1):
        add_table_row_data(risk_table, i, [risk, severity, mitigation])

    # === PAGE 4: Approvals & Authorization ===
    doc.add_page_break()

    add_heading_styled(doc, '5. Approval History', level=1)

    add_body_text(doc,
        'This request has been reviewed and endorsed by the following stakeholders:')

    approval_table = doc.add_table(rows=5, cols=4)
    approval_table.style = 'Table Grid'
    add_table_row_data(approval_table, 0, ['Reviewer', 'Title', 'Date', 'Decision'])
    for cell in approval_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    approvals = [
        ['David Chen', 'VP of Engineering', 'Feb 28, 2025', 'Approved'],
        ['Lisa Rodriguez', 'Director of IT Operations', 'Mar 3, 2025', 'Approved'],
        ['Michael Park', 'VP of Finance', 'Mar 8, 2025', 'Approved with conditions'],
        ['Angela Williams', 'Chief Operating Officer', 'Mar 12, 2025', 'Approved'],
    ]
    for i, row_data in enumerate(approvals, 1):
        add_table_row_data(approval_table, i, row_data)

    doc.add_paragraph()

    add_body_text(doc,
        'Conditions from VP of Finance: Monthly expenditure reports to be submitted '
        'to Finance by the 5th of each month. Any single purchase order exceeding '
        '$100,000 requires separate VP Finance approval.')

    doc.add_paragraph()

    add_heading_styled(doc, '6. Authorization', level=1)

    add_body_text(doc,
        'I hereby authorize the allocation of funds not to exceed $2,450,000.00 for '
        'the Enterprise Data Center Expansion Phase II project as described in this '
        'document. This authorization is subject to the conditions noted above and '
        'compliance with all applicable corporate procurement policies.')

    doc.add_paragraph()

    add_body_text(doc, 'Authorized by:', bold=True)

    # Leave empty space after "Authorized by:" - no signature line
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
