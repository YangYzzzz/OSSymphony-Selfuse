"""
Initial Setup: Protect tracked changes with password
Task ID: writer_lec_079
Domain: libreoffice_writer

Creates a Writer document with 12 tracked changes from a review cycle.
Track changes recording is active. No protection is set.
"""

import os
import shlex
import subprocess
import time
import zipfile
import shutil
import tempfile
from io import BytesIO
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_079'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# Word XML namespaces
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

NSMAP = {
    'w': W_NS,
    'r': R_NS,
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
}


def qn(tag):
    """Expand namespace prefix, e.g. 'w:body' -> '{http://...}body'."""
    ns_map = {
        'w': W_NS,
        'r': R_NS,
    }
    prefix, local = tag.split(':')
    return f'{{{ns_map[prefix]}}}{local}'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    """Build a .docx file with tracked changes by constructing XML directly."""
    from docx import Document

    # --- Step 1: Create base document with python-docx ---
    doc = Document()

    # Set document title in core properties
    doc.core_properties.title = "Quarterly Marketing Strategy Review"
    doc.core_properties.author = "Elena Rodriguez"

    # Add content paragraphs (the base document text before revisions)
    doc.add_heading("Quarterly Marketing Strategy Review", level=1)
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "This document outlines the marketing strategy for Q2 2025, "
        "including budget allocations, campaign timelines, and key performance "
        "indicators. The strategy focuses on expanding our digital presence "
        "while maintaining strong engagement across traditional channels."
    )
    doc.add_heading("Budget Overview", level=2)
    doc.add_paragraph(
        "The total marketing budget for Q2 is $485,000, representing a "
        "12% increase from the previous quarter. This allocation reflects "
        "our commitment to scaling digital advertising efforts and launching "
        "new product campaigns."
    )

    # Budget breakdown table
    table = doc.add_table(rows=6, cols=3, style='Table Grid')
    headers = ['Category', 'Q1 Actual', 'Q2 Planned']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ['Digital Advertising', '$145,000', '$175,000'],
        ['Content Marketing', '$68,000', '$82,000'],
        ['Events & Sponsorships', '$95,000', '$88,000'],
        ['Social Media', '$52,000', '$72,000'],
        ['Email Campaigns', '$35,000', '$68,000'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_heading("Campaign Timeline", level=2)
    doc.add_paragraph(
        "The Spring Product Launch campaign will run from April 7 through "
        "May 16, with a focus on social media engagement and influencer "
        "partnerships. The Summer Preview campaign begins June 1 and extends "
        "through June 30."
    )
    doc.add_heading("Key Performance Indicators", level=2)
    doc.add_paragraph(
        "We will track the following metrics to evaluate campaign success:"
    )
    doc.add_paragraph("Website traffic: Target 25% increase in unique visitors", style="List Bullet")
    doc.add_paragraph("Lead generation: Target 500 new qualified leads per month", style="List Bullet")
    doc.add_paragraph("Social media engagement: Target 18% engagement rate", style="List Bullet")
    doc.add_paragraph("Email open rate: Target 32% open rate across all campaigns", style="List Bullet")
    doc.add_paragraph("Conversion rate: Target 4.5% conversion on landing pages", style="List Bullet")

    doc.add_heading("Team Assignments", level=2)
    doc.add_paragraph(
        "The campaign execution team consists of Sarah Chen leading digital "
        "advertising, Marcus Johnson overseeing content creation, Priya Patel "
        "managing social media channels, and David Kim coordinating event logistics."
    )
    doc.add_heading("Risk Assessment", level=2)
    doc.add_paragraph(
        "Primary risks include potential budget constraints due to economic "
        "conditions, increased competition in digital ad space, and possible "
        "delays in creative asset delivery. Mitigation strategies are detailed "
        "in Appendix B."
    )
    doc.add_heading("Next Steps", level=2)
    doc.add_paragraph(
        "The team will finalize vendor contracts by March 28, complete creative "
        "briefs by April 2, and begin campaign deployment on April 7 as scheduled."
    )

    # Save the base document first
    tmp_base = tempfile.mktemp(suffix='.docx')
    doc.save(tmp_base)

    # --- Step 2: Inject tracked changes via XML manipulation ---
    # We'll manipulate the document.xml inside the zip
    tmp_dir = tempfile.mkdtemp()
    with zipfile.ZipFile(tmp_base, 'r') as zin:
        zin.extractall(tmp_dir)

    doc_xml_path = os.path.join(tmp_dir, 'word', 'document.xml')
    tree = etree.parse(doc_xml_path)
    root = tree.getroot()
    nsmap = {'w': W_NS}

    body = root.find('.//w:body', nsmap)
    paragraphs = body.findall('.//w:p', nsmap)

    # Helper to create revision elements
    rev_id_counter = [1]
    AUTHOR = "James Morrison"
    DATE = "2025-03-20T14:30:00Z"

    def next_rev_id():
        rid = rev_id_counter[0]
        rev_id_counter[0] += 1
        return str(rid)

    def make_run_props(bold=False, italic=False):
        """Create w:rPr element."""
        rpr = etree.SubElement(etree.Element('dummy'), qn('w:rPr'))
        if bold:
            etree.SubElement(rpr, qn('w:b'))
        if italic:
            etree.SubElement(rpr, qn('w:i'))
        return rpr

    def add_insertion(para, text, author=AUTHOR, date=DATE, bold=False, italic=False, before_element=None):
        """Add an insertion tracked change to a paragraph."""
        ins = etree.Element(qn('w:ins'))
        ins.set(qn('w:id'), next_rev_id())
        ins.set(qn('w:author'), author)
        ins.set(qn('w:date'), date)

        run = etree.SubElement(ins, qn('w:r'))
        if bold or italic:
            rpr = etree.SubElement(run, qn('w:rPr'))
            if bold:
                etree.SubElement(rpr, qn('w:b'))
            if italic:
                etree.SubElement(rpr, qn('w:i'))
        t = etree.SubElement(run, qn('w:t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text

        if before_element is not None:
            before_element.addprevious(ins)
        else:
            para.append(ins)
        return ins

    def add_deletion(para, text, author=AUTHOR, date=DATE, before_element=None):
        """Add a deletion tracked change to a paragraph."""
        delete = etree.Element(qn('w:del'))
        delete.set(qn('w:id'), next_rev_id())
        delete.set(qn('w:author'), author)
        delete.set(qn('w:date'), date)

        run = etree.SubElement(delete, qn('w:r'))
        dt = etree.SubElement(run, qn('w:delText'))
        dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        dt.text = text

        if before_element is not None:
            before_element.addprevious(delete)
        else:
            para.append(delete)
        return delete

    def add_format_change(para, run_elem, author=AUTHOR, date=DATE, old_bold=None, old_italic=None):
        """Add a format revision (rPrChange) to a run in a paragraph."""
        rpr = run_elem.find(qn('w:rPr'), nsmap)
        if rpr is None:
            rpr = etree.SubElement(run_elem, qn('w:rPr'))
            run_elem.insert(0, rpr)  # rPr must be first child

        rpr_change = etree.SubElement(rpr, qn('w:rPrChange'))
        rpr_change.set(qn('w:id'), next_rev_id())
        rpr_change.set(qn('w:author'), author)
        rpr_change.set(qn('w:date'), date)

        old_rpr = etree.SubElement(rpr_change, qn('w:rPr'))
        if old_bold is not None:
            b = etree.SubElement(old_rpr, qn('w:b'))
            if not old_bold:
                b.set(qn('w:val'), '0')
        if old_italic is not None:
            i_elem = etree.SubElement(old_rpr, qn('w:i'))
            if not old_italic:
                i_elem.set(qn('w:val'), '0')

    # --- Apply 12 tracked changes ---

    # Change 1: Insertion in Executive Summary (para index ~3, after heading)
    # Add text " and measurable outcomes" to the exec summary paragraph
    if len(paragraphs) > 3:
        p = paragraphs[3]
        add_insertion(p, " and measurable outcomes",
                     author="James Morrison", date="2025-03-20T10:15:00Z")

    # Change 2: Deletion in Budget Overview paragraph
    if len(paragraphs) > 5:
        p = paragraphs[5]
        add_deletion(p, "significantly ",
                    author="Aisha Thompson", date="2025-03-20T11:02:00Z")

    # Change 3: Insertion - add clarifying text about digital advertising
    if len(paragraphs) > 5:
        p = paragraphs[5]
        add_insertion(p, " across all major platforms",
                     author="James Morrison", date="2025-03-20T11:30:00Z")

    # Change 4: Deletion in campaign timeline
    # Find the campaign timeline paragraph
    timeline_idx = None
    for i, p in enumerate(paragraphs):
        runs = p.findall('.//w:r/w:t', nsmap)
        for r in runs:
            if r.text and 'Spring Product Launch' in r.text:
                timeline_idx = i
                break
        if timeline_idx is not None:
            break

    if timeline_idx is not None:
        p = paragraphs[timeline_idx]
        add_deletion(p, "preliminary ",
                    author="Priya Patel", date="2025-03-21T09:45:00Z")

    # Change 5: Insertion in campaign timeline
    if timeline_idx is not None:
        p = paragraphs[timeline_idx]
        add_insertion(p, " with dedicated tracking dashboards",
                     author="Priya Patel", date="2025-03-21T09:50:00Z")

    # Change 6: Insertion of new bullet point (as a new paragraph with tracked change)
    # Find the last bullet point
    bullet_indices = []
    for i, p in enumerate(paragraphs):
        ppr = p.find(qn('w:pPr'), nsmap)
        if ppr is not None:
            pstyle = ppr.find(qn('w:pStyle'), nsmap)
            if pstyle is not None and 'List' in pstyle.get(qn('w:val'), ''):
                bullet_indices.append(i)

    if bullet_indices:
        last_bullet = paragraphs[bullet_indices[-1]]
        # Create a new inserted paragraph after the last bullet
        new_p = etree.Element(qn('w:p'))
        # Copy paragraph properties from last bullet
        ppr_src = last_bullet.find(qn('w:pPr'), nsmap)
        if ppr_src is not None:
            new_ppr = etree.fromstring(etree.tostring(ppr_src))
            # Add paragraph revision mark
            rpr_elem = etree.SubElement(new_ppr, qn('w:rPr'))
            ins_elem = etree.SubElement(rpr_elem, qn('w:ins'))
            ins_elem.set(qn('w:id'), next_rev_id())
            ins_elem.set(qn('w:author'), "Marcus Johnson")
            ins_elem.set(qn('w:date'), "2025-03-21T14:20:00Z")
            new_p.append(new_ppr)

        ins = etree.SubElement(new_p, qn('w:ins'))
        ins.set(qn('w:id'), next_rev_id())
        ins.set(qn('w:author'), "Marcus Johnson")
        ins.set(qn('w:date'), "2025-03-21T14:20:00Z")
        run = etree.SubElement(ins, qn('w:r'))
        t = etree.SubElement(run, qn('w:t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = "Brand awareness: Target 15% increase in aided brand recall"
        last_bullet.addnext(new_p)

    # Change 7: Format change - make a run bold (tracked as formatting revision)
    # Find team assignments paragraph
    team_idx = None
    for i, p in enumerate(paragraphs):
        runs = p.findall('.//w:r/w:t', nsmap)
        for r in runs:
            if r.text and 'Sarah Chen' in r.text:
                team_idx = i
                break
        if team_idx is not None:
            break

    if team_idx is not None:
        p = paragraphs[team_idx]
        first_run = p.find(qn('w:r'), nsmap)
        if first_run is not None:
            # Add bold to the run and mark it as a format change
            rpr = first_run.find(qn('w:rPr'), nsmap)
            if rpr is None:
                rpr = etree.SubElement(first_run, qn('w:rPr'))
                first_run.insert(0, rpr)
            etree.SubElement(rpr, qn('w:b'))
            add_format_change(p, first_run,
                            author="Aisha Thompson", date="2025-03-21T15:10:00Z",
                            old_bold=False)

    # Change 8: Deletion in risk assessment
    risk_idx = None
    for i, p in enumerate(paragraphs):
        runs = p.findall('.//w:r/w:t', nsmap)
        for r in runs:
            if r.text and 'Primary risks' in r.text:
                risk_idx = i
                break
        if risk_idx is not None:
            break

    if risk_idx is not None:
        p = paragraphs[risk_idx]
        add_deletion(p, "possible ",
                    author="David Kim", date="2025-03-22T08:30:00Z")

    # Change 9: Insertion in risk assessment
    if risk_idx is not None:
        p = paragraphs[risk_idx]
        add_insertion(p, " and contingency plans have been approved by leadership",
                     author="David Kim", date="2025-03-22T08:35:00Z")

    # Change 10: Deletion in next steps
    next_idx = None
    for i, p in enumerate(paragraphs):
        runs = p.findall('.//w:r/w:t', nsmap)
        for r in runs:
            if r.text and 'finalize vendor' in r.text:
                next_idx = i
                break
        if next_idx is not None:
            break

    if next_idx is not None:
        p = paragraphs[next_idx]
        add_deletion(p, "as scheduled",
                    author="Elena Rodriguez", date="2025-03-22T10:00:00Z")

    # Change 11: Insertion in next steps
    if next_idx is not None:
        p = paragraphs[next_idx]
        add_insertion(p, "pending final budget approval from the CFO",
                     author="Elena Rodriguez", date="2025-03-22T10:05:00Z")

    # Change 12: Format change - italicize in executive summary
    if len(paragraphs) > 3:
        p = paragraphs[3]
        first_run = p.find(qn('w:r'), nsmap)
        if first_run is not None:
            rpr = first_run.find(qn('w:rPr'), nsmap)
            if rpr is None:
                rpr = etree.SubElement(first_run, qn('w:rPr'))
                first_run.insert(0, rpr)
            etree.SubElement(rpr, qn('w:i'))
            add_format_change(p, first_run,
                            author="Aisha Thompson", date="2025-03-22T11:00:00Z",
                            old_italic=False)

    # --- Step 3: Enable track changes recording in settings ---
    settings_path = os.path.join(tmp_dir, 'word', 'settings.xml')
    if os.path.exists(settings_path):
        settings_tree = etree.parse(settings_path)
        settings_root = settings_tree.getroot()
    else:
        # Create minimal settings.xml
        settings_root = etree.Element(qn('w:settings'))
        settings_root.set('xmlns:w', W_NS)
        settings_tree = etree.ElementTree(settings_root)

    # Add trackRevisions element to enable track changes
    # Remove existing if present
    for elem in settings_root.findall(qn('w:trackRevisions'), nsmap):
        settings_root.remove(elem)
    track_rev = etree.SubElement(settings_root, qn('w:trackRevisions'))

    # Save settings
    settings_tree.write(settings_path, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Save modified document.xml
    tree.write(doc_xml_path, xml_declaration=True, encoding='UTF-8', standalone=True)

    # --- Step 4: Repack the docx ---
    with zipfile.ZipFile(OUTPUT, 'w', zipfile.ZIP_DEFLATED) as zout:
        for dirpath, dirnames, filenames in os.walk(tmp_dir):
            for fn in filenames:
                full_path = os.path.join(dirpath, fn)
                arcname = os.path.relpath(full_path, tmp_dir)
                zout.write(full_path, arcname)

    # Cleanup
    os.remove(tmp_base)
    shutil.rmtree(tmp_dir)
    print(f'Initial file created: {OUTPUT}')

    # --- Step 5: GUI-ready startup ---
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
