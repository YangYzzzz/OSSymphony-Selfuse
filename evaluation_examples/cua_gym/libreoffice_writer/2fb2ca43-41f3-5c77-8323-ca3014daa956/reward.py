"""
Reward Script: Create a list style named 'LegalNumbering' using Article I/II/III format
Task ID: writer_list_047
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): 'LegalNumbering' paragraph style exists in the document
  Component 2 (0.3): The associated numbering uses upperRoman format and 'Article %1' level text
  Component 3 (0.3): All 5 paragraphs have the 'LegalNumbering' style applied
"""

import os
import zipfile
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_list_047'

FILE_PATH = f'{WORKDIR}/Desktop/contract_sections.docx'


def verify_task(file_path):
    """
    Verify that a 'LegalNumbering' list style was created with uppercase Roman
    numeral format and 'Article ' prefix, and applied to all five paragraphs.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must exist and be a valid docx
    if not os.path.exists(file_path):
        print(f"CRITICAL: File not found: {file_path}")
        print("REWARD: 0.0")
        return 0.0

    try:
        zf = zipfile.ZipFile(file_path, 'r')
    except Exception as e:
        print(f"CRITICAL: Cannot open docx as zip: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ------------------------------------------------------------------
    # Component 1: 'LegalNumbering' paragraph style exists (0.4 points)
    # ------------------------------------------------------------------
    # This FAILS on initial_env (style does not exist) and PASSES on golden_env
    try:
        styles_xml = zf.read('word/styles.xml').decode('utf-8', errors='replace')
        # Look for a style named 'LegalNumbering' (by name attribute or styleId)
        has_legal_style = (
            'LegalNumbering' in styles_xml
            and re.search(r'w:type="paragraph"[^>]*LegalNumbering|LegalNumbering[^>]*w:type="paragraph"', styles_xml) is not None
        )
        # Also accept: the styleId alone is sufficient (name is implicitly same)
        if not has_legal_style:
            has_legal_style = bool(re.search(
                r'<w:style[^>]*w:styleId="LegalNumbering"', styles_xml
            ))
        if has_legal_style:
            print("PASS: Component 1 — 'LegalNumbering' paragraph style exists in styles.xml (0.4 pts)")
            total_score += 0.4
        else:
            print("FAIL: Component 1 — 'LegalNumbering' paragraph style NOT found in styles.xml")
    except Exception as e:
        print(f"ERROR: Component 1 — could not read styles.xml: {e}")

    # ------------------------------------------------------------------
    # Component 2: Numbering definition uses upperRoman + 'Article %1' (0.3 points)
    # ------------------------------------------------------------------
    # The LegalNumbering style must be linked to a numbering abstractNum that has:
    #   numFmt = upperRoman  AND  lvlText containing 'Article'
    # This distinguishes it from any generic list style; only the golden_env satisfies this.
    try:
        if 'word/numbering.xml' in zf.namelist():
            numbering_xml = zf.read('word/numbering.xml').decode('utf-8', errors='replace')

            # First find the abstractNum associated with 'LegalNumbering' by its name attribute
            abstract_match = re.search(
                r'<w:abstractNum\s[^>]*>(.*?)</w:abstractNum>',
                numbering_xml,
                re.DOTALL
            )
            # Find the abstract num block that contains the name LegalNumbering
            legal_abstract_block = None
            for block in re.finditer(
                r'<w:abstractNum\s[^>]*>(.*?)</w:abstractNum>',
                numbering_xml,
                re.DOTALL
            ):
                if 'LegalNumbering' in block.group(0):
                    legal_abstract_block = block.group(0)
                    break

            if legal_abstract_block is None:
                print("FAIL: Component 2 — No abstractNum with name 'LegalNumbering' found in numbering.xml")
            else:
                has_upper_roman = bool(re.search(r'w:val="upperRoman"', legal_abstract_block))
                # lvlText should be 'Article %1' (with trailing space before %1)
                lvl_text_match = re.search(r'<w:lvlText\s+w:val="([^"]*)"', legal_abstract_block)
                lvl_text_val = lvl_text_match.group(1) if lvl_text_match else ''
                has_article_prefix = 'Article' in lvl_text_val and '%1' in lvl_text_val

                if has_upper_roman and has_article_prefix:
                    print(f"PASS: Component 2 — Numbering uses upperRoman format with lvlText='{lvl_text_val}' (0.3 pts)")
                    total_score += 0.3
                else:
                    print(
                        f"FAIL: Component 2 — Expected upperRoman=True and lvlText containing 'Article %1'. "
                        f"Found: upperRoman={has_upper_roman}, lvlText='{lvl_text_val}'"
                    )
        else:
            print("FAIL: Component 2 — numbering.xml not present in docx")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ------------------------------------------------------------------
    # Component 3: All 5 paragraphs use the 'LegalNumbering' style (0.3 points)
    # ------------------------------------------------------------------
    # Only the golden_env has all 5 paragraphs assigned to 'LegalNumbering'
    try:
        from docx import Document
        doc = Document(file_path)

        total_paras = len(doc.paragraphs)
        legal_paras = [p for p in doc.paragraphs if p.style.name == 'LegalNumbering']
        count_legal = len(legal_paras)

        expected_count = 5
        if count_legal == expected_count:
            print(f"PASS: Component 3 — All {expected_count} paragraphs use 'LegalNumbering' style (0.3 pts)")
            total_score += 0.3
        elif count_legal > 0:
            print(
                f"FAIL: Component 3 — Only {count_legal} of {expected_count} paragraphs use 'LegalNumbering' style"
            )
        else:
            print(f"FAIL: Component 3 — No paragraphs use 'LegalNumbering' style (total paragraphs: {total_paras})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 4)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
