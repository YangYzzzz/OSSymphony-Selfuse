"""
Initial Setup: Apply 'Keep with next paragraph' to all figure captions
Task ID: writer_tech_086
Domain: libreoffice_writer

Creates a technical report document with multiple figures and captions.
Some captions are positioned such that page breaks can separate them from images.
Captions do NOT have keep_with_next enabled.
"""

import os
import shlex
import subprocess
import time
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image as PILImage

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_086'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_placeholder_image(path, width=400, height=250, color=(70, 130, 180), label="Figure"):
    """Create a simple placeholder image with a label."""
    img = PILImage.new('RGB', (width, height), color)
    img.save(path)


def add_figure_caption(doc, caption_text, bold_label=True):
    """Add a figure caption paragraph (no keep_with_next)."""
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    # Explicitly ensure keep_with_next is NOT set
    cap.paragraph_format.keep_with_next = False
    if bold_label:
        parts = caption_text.split(":", 1)
        if len(parts) == 2:
            run_label = cap.add_run(parts[0] + ":")
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_label.font.name = "Liberation Sans"
            run_desc = cap.add_run(parts[1])
            run_desc.font.size = Pt(10)
            run_desc.font.name = "Liberation Sans"
        else:
            run = cap.add_run(caption_text)
            run.font.size = Pt(10)
            run.font.name = "Liberation Sans"
    else:
        run = cap.add_run(caption_text)
        run.font.size = Pt(10)
        run.font.name = "Liberation Sans"
    return cap


def add_body_text(doc, text, space_after=Pt(8)):
    """Add a body paragraph with consistent formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"
    return p


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading("Network Performance Analysis Report", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.paragraph_format.space_after = Pt(20)
    r = meta.add_run("Prepared by: Dr. Elena Rodriguez, Network Engineering Division\nDate: March 2026")
    r.font.size = Pt(11)
    r.font.name = "Liberation Sans"
    r.font.color.rgb = RGBColor(100, 100, 100)

    # --- Section 1 ---
    doc.add_heading("1. Executive Summary", level=1)
    add_body_text(doc,
        "This report presents a comprehensive analysis of our enterprise network "
        "performance over the first quarter of 2026. The study covers latency measurements, "
        "throughput benchmarks, packet loss statistics, and bandwidth utilization across "
        "all three regional data centers. Key findings indicate a 23% improvement in "
        "average response times following the infrastructure upgrade completed in January.")

    add_body_text(doc,
        "The analysis methodology employed continuous monitoring using NetFlow collectors "
        "deployed at strategic egress points, supplemented by active probing from 47 "
        "geographically distributed test endpoints. Statistical significance was verified "
        "using two-tailed t-tests with a confidence level of 95%.")

    # --- Section 2 ---
    doc.add_heading("2. Latency Analysis", level=1)
    add_body_text(doc,
        "Round-trip latency measurements were collected at five-minute intervals from "
        "each regional endpoint to the primary data center in Portland. The following "
        "figure shows the distribution of latency values before and after the upgrade.")

    add_body_text(doc,
        "Pre-upgrade measurements showed a median latency of 45ms with a 95th percentile "
        "of 127ms. Post-upgrade values improved to a median of 34ms with a 95th percentile "
        "of 78ms. The improvement was most pronounced during peak traffic hours between "
        "09:00 and 17:00 UTC, where the median dropped from 62ms to 38ms.")

    # Figure 1 - image placeholder
    img1_path = f'{WORKDIR}/fig_latency.png'
    create_placeholder_image(img1_path, 480, 300, (41, 98, 163), "Latency Distribution")
    p_img1 = doc.add_paragraph()
    p_img1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img1.add_run().add_picture(img1_path, width=Inches(4.5))

    # Caption 1
    add_figure_caption(doc, "Figure 1: Latency distribution comparison — pre-upgrade (blue) vs post-upgrade (green) across all regional endpoints")

    add_body_text(doc,
        "The latency improvements were consistent across all three regions, though the "
        "Asia-Pacific endpoints showed the most dramatic gains. This is attributable to "
        "the new peering arrangements established with transit providers in Singapore "
        "and Tokyo during the upgrade window.")

    add_body_text(doc,
        "It is worth noting that the variance in latency also decreased significantly. "
        "The interquartile range narrowed from 31ms to 18ms, suggesting more predictable "
        "network behavior that benefits real-time applications and video conferencing.")

    # --- Section 3 ---
    doc.add_heading("3. Throughput Benchmarks", level=1)
    add_body_text(doc,
        "Throughput testing was conducted using iperf3 sessions running for 60 seconds "
        "each, with parallel streams set to 8 to simulate realistic multi-connection "
        "workloads. Tests were run between each regional pair at four-hour intervals.")

    add_body_text(doc,
        "The Portland-to-Chicago link showed the highest improvement, with sustained "
        "throughput increasing from 7.2 Gbps to 9.4 Gbps. The Portland-to-Frankfurt "
        "transatlantic link improved from 3.1 Gbps to 4.8 Gbps, primarily due to the "
        "activation of the new submarine cable path via the Hibernia Express route.")

    add_body_text(doc,
        "Aggregate throughput across all monitored links reached a peak of 42.7 Gbps "
        "during the February traffic surge associated with the product launch event. "
        "No congestion-related packet drops were observed during this period, validating "
        "the capacity planning model used in the upgrade design phase.")

    # Figure 2
    img2_path = f'{WORKDIR}/fig_throughput.png'
    create_placeholder_image(img2_path, 480, 280, (34, 139, 34), "Throughput Over Time")
    p_img2 = doc.add_paragraph()
    p_img2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img2.add_run().add_picture(img2_path, width=Inches(4.5))

    # Caption 2
    add_figure_caption(doc, "Figure 2: Sustained throughput measurements across regional links, January through March 2026")

    add_body_text(doc,
        "The throughput gains were particularly evident in TCP connections with window "
        "scaling enabled. The new router firmware includes an optimized congestion control "
        "algorithm (BBRv3) that better utilizes available bandwidth on high-latency paths.")

    # --- Section 4 ---
    doc.add_heading("4. Packet Loss and Error Rates", level=1)
    add_body_text(doc,
        "Packet loss monitoring was performed using ICMP echo probes sent at one-second "
        "intervals. The loss rate is calculated as the percentage of probes that either "
        "timed out (threshold: 3 seconds) or returned with an error code.")

    add_body_text(doc,
        "Prior to the upgrade, the aggregate packet loss rate averaged 0.34% with "
        "occasional spikes reaching 2.1% during peak hours. After the upgrade, the "
        "average loss rate dropped to 0.08%, with a maximum observed spike of 0.52%. "
        "These improvements directly correlate with the replacement of aging Juniper MX240 "
        "routers with Cisco 8201-32FH units at the core aggregation layer.")

    # Figure 3
    img3_path = f'{WORKDIR}/fig_packetloss.png'
    create_placeholder_image(img3_path, 480, 260, (178, 34, 34), "Packet Loss Timeline")
    p_img3 = doc.add_paragraph()
    p_img3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img3.add_run().add_picture(img3_path, width=Inches(4.5))

    # Caption 3
    add_figure_caption(doc, "Figure 3: Daily packet loss rate over Q1 2026, showing the upgrade cutover on January 18th")

    add_body_text(doc,
        "CRC error rates on optical interfaces also dropped to near zero after the "
        "transceiver modules were replaced as part of the upgrade. Previously, two "
        "100GbE links on the Portland-Chicago route were experiencing intermittent "
        "CRC errors averaging 12 per hour, attributed to degraded QSFP28 modules.")

    # --- Section 5 ---
    doc.add_heading("5. Bandwidth Utilization Patterns", level=1)
    add_body_text(doc,
        "Bandwidth utilization was tracked using SNMP polling of interface counters at "
        "five-minute granularity. The data was aggregated per-link and per-region to "
        "identify utilization trends and capacity planning thresholds.")

    add_body_text(doc,
        "The average link utilization across all WAN connections settled at 47% post-upgrade, "
        "down from 68% pre-upgrade. This headroom provides comfortable margin for traffic "
        "growth projections through Q4 2026 without additional capacity investment. The "
        "busiest single link (Portland to Chicago) peaked at 72% utilization during the "
        "product launch event but averaged 53% over the quarter.")

    # Figure 4
    img4_path = f'{WORKDIR}/fig_bandwidth.png'
    create_placeholder_image(img4_path, 480, 280, (148, 103, 189), "Bandwidth Utilization")
    p_img4 = doc.add_paragraph()
    p_img4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img4.add_run().add_picture(img4_path, width=Inches(4.5))

    # Caption 4
    add_figure_caption(doc, "Figure 4: Bandwidth utilization heatmap by region and time-of-day for Q1 2026")

    add_body_text(doc,
        "An interesting finding was the shift in traffic patterns to more uniform "
        "distribution across time zones. The deployment of CDN edge nodes in the "
        "Singapore and Frankfurt data centers redirected approximately 15% of traffic "
        "that previously traversed the transatlantic links during European business hours.")

    # --- Section 6 ---
    doc.add_heading("6. Quality of Service Metrics", level=1)
    add_body_text(doc,
        "QoS measurements focused on three key indicators relevant to real-time "
        "communications: jitter (variation in latency), mean opinion score (MOS) for "
        "voice traffic, and video stream quality metrics including resolution stability "
        "and buffering frequency.")

    add_body_text(doc,
        "Voice traffic jitter decreased from an average of 8.3ms to 3.7ms post-upgrade, "
        "bringing all monitored paths within the ITU-T G.114 recommended threshold of "
        "30ms one-way delay. The corresponding MOS scores improved from 3.8 to 4.3 on "
        "the five-point scale, representing a noticeable quality improvement for end users.")

    # Figure 5
    img5_path = f'{WORKDIR}/fig_qos.png'
    create_placeholder_image(img5_path, 480, 300, (255, 152, 0), "QoS Dashboard")
    p_img5 = doc.add_paragraph()
    p_img5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img5.add_run().add_picture(img5_path, width=Inches(4.5))

    # Caption 5
    add_figure_caption(doc, "Figure 5: Quality of service metrics dashboard — voice MOS, video resolution stability, and jitter measurements")

    # --- Section 7 ---
    doc.add_heading("7. Conclusions and Recommendations", level=1)
    add_body_text(doc,
        "The Q1 2026 infrastructure upgrade has delivered measurable improvements across "
        "all monitored network performance dimensions. Latency, throughput, packet loss, "
        "and QoS metrics all show statistically significant gains that align with or "
        "exceed the targets established in the project planning phase.")

    add_body_text(doc,
        "Based on these findings, we recommend the following next steps: (1) Extend the "
        "BBRv3 congestion control deployment to the remaining edge routers in the APAC "
        "region by Q2 2026. (2) Implement automated anomaly detection using the baseline "
        "metrics established in this report. (3) Begin planning for 400GbE backbone "
        "upgrades on the Portland-Chicago route to accommodate projected growth beyond "
        "Q4 2026.")

    add_body_text(doc,
        "The monitoring infrastructure itself should be enhanced with real-time streaming "
        "telemetry (gNMI/gRPC) to replace SNMP polling, enabling sub-second visibility "
        "into network state changes and faster incident response times.")

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Clean up temp images
    for img_file in ['fig_latency.png', 'fig_throughput.png', 'fig_packetloss.png',
                     'fig_bandwidth.png', 'fig_qos.png']:
        img_path = f'{WORKDIR}/{img_file}'
        if os.path.exists(img_path):
            os.remove(img_path)

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
