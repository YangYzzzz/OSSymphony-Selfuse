"""
Initial Setup: Create event_plan.docx with ten plain paragraphs (no list styles)
Task ID: writer_list_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document

WORKDIR = '/home/user/Desktop'  # VM path — task context says ~/Desktop/event_plan.docx
TASK_ID = 'event_plan'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Ten plain paragraphs forming a two-level structure — ALL plain text, no list styles
    # Level 1: Venue Setup
    doc.add_paragraph("Venue Setup")
    doc.add_paragraph("Arrange tables and chairs")
    doc.add_paragraph("Set up audio-visual equipment")
    doc.add_paragraph("Place signage and banners")

    # Level 1: Catering
    doc.add_paragraph("Catering")
    doc.add_paragraph("Confirm menu with caterer")
    doc.add_paragraph("Arrange dietary accommodations")

    # Level 1: Guest Management
    doc.add_paragraph("Guest Management")
    doc.add_paragraph("Print name badges")
    doc.add_paragraph("Prepare welcome packets")
    doc.add_paragraph("Brief registration volunteers")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
