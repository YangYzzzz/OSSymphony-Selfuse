"""
Reward Script: Create a nested list where main items are numbered and sub-items are bulleted
Task ID: writer_list_020
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5 pts): All 3 level-1 paragraphs use a numbered list style ('List Number')
  Component 2 (0.5 pts): All 8 level-2 paragraphs use an indented bullet list style ('List Bullet 2')
  Text content preservation is a precondition gate (does not contribute to score).
"""

import os
from docx import Document

WORKDIR = '/home/user/Desktop'
FILE_PATH = f'{WORKDIR}/event_plan.docx'

# Ground truth: level-1 texts (should be numbered list)
LEVEL1_TEXTS = ['Venue Setup', 'Catering', 'Guest Management']

# Ground truth: level-2 texts (should be bullet sub-list, indented)
LEVEL2_TEXTS = [
    'Arrange tables and chairs',
    'Set up audio-visual equipment',
    'Place signage and banners',
    'Confirm menu with caterer',
    'Arrange dietary accommodations',
    'Print name badges',
    'Prepare welcome packets',
    'Brief registration volunteers',
]

# Expected full ordered list of paragraph texts
ALL_TEXTS = [
    'Venue Setup',
    'Arrange tables and chairs',
    'Set up audio-visual equipment',
    'Place signage and banners',
    'Catering',
    'Confirm menu with caterer',
    'Arrange dietary accommodations',
    'Guest Management',
    'Print name badges',
    'Prepare welcome packets',
    'Brief registration volunteers',
]


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition gate: paragraph text content must be intact
    # This does NOT contribute to score — it only gates further checking.
    actual_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if actual_texts != ALL_TEXTS:
        print(f"PRECONDITION FAIL: text content mismatch (file may be corrupt or wrong file).")
        print(f"  Expected: {ALL_TEXTS}")
        print(f"  Got:      {actual_texts}")
        print("REWARD: 0.0")
        return 0.0
    else:
        print("PRECONDITION PASS: all 11 paragraphs present with correct text.")

    # Build lookup: text -> style name
    para_style_map = {p.text.strip(): (p.style.name if p.style else '') for p in doc.paragraphs if p.text.strip()}

    # Component 1: Level-1 paragraphs (Venue Setup, Catering, Guest Management)
    # must use a numbered list style. In the initial file all are 'Normal' — so this
    # FAILS on initial and PASSES on golden. Accept 'List Number', 'List Number 2', etc.
    try:
        numbered_count = 0
        numbered_found = []
        numbered_missing = []
        for text in LEVEL1_TEXTS:
            style = para_style_map.get(text, '')
            if 'List Number' in style:
                numbered_count += 1
                numbered_found.append(f"{text!r}:{style!r}")
            else:
                numbered_missing.append(f"{text!r}:style={style!r}")

        if numbered_count == len(LEVEL1_TEXTS):
            print(f"PASS: Component 1 — all {len(LEVEL1_TEXTS)} level-1 items have numbered list style: {numbered_found} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — only {numbered_count}/{len(LEVEL1_TEXTS)} level-1 items are numbered."
                  f" Missing: {numbered_missing}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Level-2 paragraphs must use an indented bullet list style.
    # The task requires they be "indented underneath" their parent numbered item.
    # 'List Bullet 2' is the correct indented sub-bullet style in python-docx.
    # Accept 'List Bullet 2', 'List Bullet 3', etc. — any multi-level bullet style
    # that is NOT 'List Bullet' (which is top-level, non-indented).
    # In the initial file all are 'Normal' — so this FAILS on initial and PASSES on golden.
    try:
        bullet_count = 0
        bullet_found = []
        bullet_missing = []
        for text in LEVEL2_TEXTS:
            style = para_style_map.get(text, '')
            # 'List Bullet 2', 'List Bullet 3', etc. — indented bullet styles
            if 'List Bullet' in style and style != 'List Bullet':
                bullet_count += 1
                bullet_found.append(f"{text!r}:{style!r}")
            else:
                bullet_missing.append(f"{text!r}:style={style!r}")

        if bullet_count == len(LEVEL2_TEXTS):
            print(f"PASS: Component 2 — all {len(LEVEL2_TEXTS)} level-2 items have indented bullet style: {bullet_found} (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 2 — only {bullet_count}/{len(LEVEL2_TEXTS)} level-2 items have indented bullet style."
                  f" Missing: {bullet_missing}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
