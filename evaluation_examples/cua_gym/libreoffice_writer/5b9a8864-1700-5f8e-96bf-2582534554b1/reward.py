"""
Reward Script: Disciplinary Action Tracking Form
Task ID: writer_hr_068
Domain: libreoffice_writer
Scoring:
  Component 1 (0.20): Five section headings present
  Component 2 (0.25): Disciplinary Progression table (5 rows x 4 cols, correct headers/levels)
  Component 3 (0.25): SMART Goals table (6 rows x 6 cols, correct headers)
  Component 4 (0.15): Follow-Up Schedule table (>=5 rows x 4 cols)
  Component 5 (0.15): Four signature blocks (Employee, Manager/Supervisor, HR, Witness)
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_068'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect all paragraph texts for analysis
    all_para_texts = [p.text.strip() for p in doc.paragraphs]
    all_para_lower = [t.lower() for t in all_para_texts]

    # =========================================================================
    # Component 1: Five section headings present (0.20 points)
    # Task requires 5 sections. The golden file uses Heading 2 style for sections.
    # We check for heading-style paragraphs containing section keywords.
    # Initial file has 0 headings -> FAIL on initial, PASS on golden.
    # =========================================================================
    try:
        # Look for section headings (any heading style or bold text with section keywords)
        section_keywords = [
            "employee information",    # Section 1
            "disciplinary action progression",  # Section 2 (or just "progression")
            "smart goals",             # Section 3 (or "improvement plan")
            "follow-up",              # Section 4 (or "follow up")
            "signature",              # Section 5 (or "acknowledgment")
        ]

        # Also accept alternate phrasings
        alt_keywords = [
            ["employee info", "incident summary", "section 1"],
            ["progression", "disciplinary", "section 2"],
            ["smart", "improvement plan", "section 3"],
            ["follow-up", "follow up", "schedule", "section 4"],
            ["signature", "acknowledgment", "acknowledgement", "section 5"],
        ]

        heading_paras = []
        for p in doc.paragraphs:
            style_name = p.style.name.lower() if p.style else ""
            is_heading = "heading" in style_name or "title" in style_name
            # Also consider bold paragraphs as potential section headers
            is_bold_header = False
            if p.runs and all(r.font.bold for r in p.runs if r.text.strip()):
                is_bold_header = True
            if is_heading or is_bold_header:
                heading_paras.append(p.text.strip().lower())

        sections_found = 0
        for i, alt_list in enumerate(alt_keywords):
            found = False
            for heading_text in heading_paras:
                for kw in alt_list:
                    if kw in heading_text:
                        found = True
                        break
                if found:
                    break
            if found:
                sections_found += 1

        if sections_found >= 5:
            print(f"PASS: Component 1 — All 5 section headings found (0.20 pts)")
            total_score += 0.20
        elif sections_found >= 3:
            partial = round(0.20 * sections_found / 5, 2)
            print(f"PARTIAL: Component 1 — {sections_found}/5 section headings found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — Only {sections_found}/5 section headings found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # =========================================================================
    # Component 2: Disciplinary Progression table (0.25 points)
    # Must have a table with 4 action levels: Verbal Warning, Written Warning,
    # Final Warning, Termination. Table should have 4 columns (action level,
    # date, description, witness). Header + 4 data rows = 5 rows.
    # Initial file has 0 tables -> FAIL on initial.
    # =========================================================================
    try:
        progression_found = False
        progression_score = 0.0

        required_levels = ["verbal warning", "written warning", "final warning", "termination"]

        for table in doc.tables:
            # Check if this table contains the progression levels
            all_cell_texts = []
            for row in table.rows:
                for cell in row.cells:
                    all_cell_texts.append(cell.text.strip().lower())

            levels_found = sum(1 for level in required_levels if any(level in ct for ct in all_cell_texts))

            if levels_found >= 3:
                progression_found = True
                # Check row count (header + 4 levels = 5 rows)
                if len(table.rows) >= 5:
                    progression_score += 0.10
                # Check column count (at least 4: level, date, description, witness)
                if len(table.columns) >= 4:
                    progression_score += 0.05
                # Check all 4 levels present
                if levels_found == 4:
                    progression_score += 0.10
                else:
                    progression_score += round(0.10 * levels_found / 4, 2)
                break

        if progression_found:
            progression_score = min(progression_score, 0.25)
            print(f"PASS: Component 2 — Progression table found with {levels_found}/4 levels ({progression_score} pts)")
            total_score += progression_score
        else:
            print(f"FAIL: Component 2 — No disciplinary progression table found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # =========================================================================
    # Component 3: SMART Goals table (0.25 points)
    # Must have a table with 6 columns: Goal#, Specific, Measurable, Achievable,
    # Relevant, Time-bound. Should have header + 5 goal rows = 6 rows.
    # Initial file has 0 tables -> FAIL on initial.
    # =========================================================================
    try:
        smart_found = False
        smart_score = 0.0

        smart_headers = ["specific", "measurable", "achievable", "relevant", "time-bound"]
        # Also accept "time bound" without hyphen
        smart_headers_alt = ["specific", "measurable", "achievable", "relevant", "time"]

        for table in doc.tables:
            # Check first row (header) for SMART column names
            if len(table.rows) < 2:
                continue
            header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
            header_combined = " ".join(header_texts)

            smart_matches = sum(1 for h in smart_headers_alt if any(h in ht for ht in header_texts))

            if smart_matches >= 4:
                smart_found = True
                # Check column count (6: Goal#, S, M, A, R, T)
                if len(table.columns) >= 6:
                    smart_score += 0.10
                elif len(table.columns) >= 5:
                    smart_score += 0.05
                # Check row count (header + 5 goals = 6 rows)
                if len(table.rows) >= 6:
                    smart_score += 0.10
                elif len(table.rows) >= 4:
                    smart_score += 0.05
                # Check all 5 SMART headers present
                if smart_matches >= 5:
                    smart_score += 0.05
                break

        if smart_found:
            smart_score = min(smart_score, 0.25)
            print(f"PASS: Component 3 — SMART Goals table found ({smart_score} pts)")
            total_score += smart_score
        else:
            print(f"FAIL: Component 3 — No SMART goals table found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # =========================================================================
    # Component 4: Follow-Up Schedule table (0.15 points)
    # Must have a table with at least 4 review dates + header row.
    # Should contain terms like "review", "scheduled", or date-related content.
    # Initial file has 0 tables -> FAIL on initial.
    # =========================================================================
    try:
        followup_found = False
        followup_score = 0.0

        review_keywords = ["review", "scheduled", "week", "day", "month"]

        for table in doc.tables:
            # Skip tables already identified (progression has "warning", smart has "measurable")
            all_cell_texts = []
            for row in table.rows:
                for cell in row.cells:
                    all_cell_texts.append(cell.text.strip().lower())

            combined = " ".join(all_cell_texts)

            # This table likely has "review" in it and should NOT have "warning" or "measurable"
            has_review = any(kw in combined for kw in review_keywords)
            has_warning = "warning" in combined and "termination" in combined
            has_smart = "measurable" in combined and "achievable" in combined

            if has_review and not has_warning and not has_smart:
                followup_found = True
                # Check at least 4 data rows (+ header = 5 rows)
                data_rows = len(table.rows) - 1  # subtract header
                if data_rows >= 4:
                    followup_score += 0.10
                elif data_rows >= 2:
                    followup_score += 0.05
                # Check at least 3 columns
                if len(table.columns) >= 3:
                    followup_score += 0.05
                break

        if followup_found:
            followup_score = min(followup_score, 0.15)
            print(f"PASS: Component 4 — Follow-up schedule table found ({followup_score} pts)")
            total_score += followup_score
        else:
            print(f"FAIL: Component 4 — No follow-up schedule table found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # =========================================================================
    # Component 5: Four signature blocks (0.15 points)
    # Must have signature areas for: Employee, Manager/Supervisor, HR Rep, Witness.
    # Initial file has only 1 paragraph (title) -> FAIL on initial.
    # =========================================================================
    try:
        sig_keywords = {
            "employee": False,
            "manager": False,
            "supervisor": False,
            "hr representative": False,
            "hr rep": False,
            "witness": False,
        }
        # Group: employee, manager/supervisor, HR, witness
        sig_groups = {
            "employee": False,
            "manager_supervisor": False,
            "hr": False,
            "witness": False,
        }

        # Look for signature-related text in paragraphs
        # A signature block typically has "Signature:" and a role name nearby
        for i, text in enumerate(all_para_lower):
            if "employee" in text and ("signature" in text or
                (i + 1 < len(all_para_lower) and "signature" in all_para_lower[i + 1])):
                sig_groups["employee"] = True
            if ("manager" in text or "supervisor" in text) and ("signature" in text or
                (i + 1 < len(all_para_lower) and "signature" in all_para_lower[i + 1])):
                sig_groups["manager_supervisor"] = True
            if ("hr" in text or "human resource" in text) and ("signature" in text or
                (i + 1 < len(all_para_lower) and "signature" in all_para_lower[i + 1])):
                sig_groups["hr"] = True
            if "witness" in text and ("signature" in text or
                (i + 1 < len(all_para_lower) and "signature" in all_para_lower[i + 1])):
                sig_groups["witness"] = True

        sigs_found = sum(1 for v in sig_groups.values() if v)

        if sigs_found >= 4:
            print(f"PASS: Component 5 — All 4 signature blocks found (0.15 pts)")
            total_score += 0.15
        elif sigs_found >= 2:
            partial = round(0.15 * sigs_found / 4, 2)
            print(f"PARTIAL: Component 5 — {sigs_found}/4 signature blocks found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 — Only {sigs_found}/4 signature blocks found")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    final_score = round(final_score, 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
