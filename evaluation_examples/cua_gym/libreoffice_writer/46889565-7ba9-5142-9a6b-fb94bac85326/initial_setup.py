"""
Initial Setup: User Conference Recap Document with Photo Placeholder Text
Task ID: writer_mktg_048
Domain: libreoffice_writer

Creates a 4-page annual user conference recap document with plain text
photo markers at paragraphs 3, 7, and 10 that need to be replaced
with properly framed image placeholders.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
DESKTOP_PATH = f'{WORKDIR}/Desktop/user_conference_recap.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ---- Title ----
    title = doc.add_heading('Annual User Conference 2025: A Recap', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Connecting, Learning, and Innovating Together')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.italic = True
        run.font.size = Pt(13)

    doc.add_paragraph()  # blank line

    # ---- Paragraph 1: Introduction ----
    p1 = doc.add_paragraph(
        'The 2025 Annual User Conference was held from June 12-14 at the Grand Metropolitan '
        'Convention Center in Chicago, Illinois. This year\'s theme, "Building Tomorrow Together," '
        'brought together over 3,400 attendees from 42 countries to share insights, explore new '
        'product capabilities, and forge lasting professional connections. The three-day event '
        'featured more than 120 sessions, workshops, and hands-on labs designed to help users '
        'maximize their investment in our platform.'
    )

    # ---- Paragraph 2: Opening Day ----
    doc.add_heading('Day One: Setting the Stage', level=2)
    p2 = doc.add_paragraph(
        'Conference registration opened at 7:00 AM on June 12, and by the time doors to the main '
        'hall opened, a line had already formed stretching around the convention center\'s east '
        'wing. Attendees collected their badges, conference materials, and the always-popular tote '
        'bags before heading to networking breakfasts organized by industry vertical. The energy '
        'was palpable as colleagues who had only spoken online finally met face-to-face, and '
        'newcomers discovered the welcoming community they had joined.'
    )

    # ---- Paragraph 3: Keynote paragraph → PHOTO MARKER HERE ----
    p3 = doc.add_paragraph(
        'The morning kicked off with a standing-room-only opening keynote in the 2,000-seat Main '
        'Auditorium. CEO Patricia Walcott took the stage to thunderous applause, delivering an '
        'energizing address that outlined the product roadmap for the next 18 months.'
    )

    # Photo marker 1 — plain text placeholder (to be replaced by agent)
    photo1 = doc.add_paragraph('[Photo: Keynote]')
    photo1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Paragraph 4: Keynote continued ----
    p4 = doc.add_paragraph(
        'Walcott\'s address highlighted three core pillars of the upcoming release: enhanced AI '
        'automation, deeper third-party integrations, and a completely redesigned mobile '
        'experience. She was joined on stage by CTO Dr. Rajesh Patel, who gave a live demonstration '
        'of the new intelligent workflow engine that can reduce manual data entry by up to 67%. '
        'The audience responded enthusiastically, with multiple standing ovations throughout the '
        '45-minute presentation.'
    )

    # ---- Paragraph 5: Expo Hall intro ----
    doc.add_heading('The Exhibition Hall Experience', level=2)
    p5 = doc.add_paragraph(
        'Following the keynote, attendees flooded the 40,000 square-foot exhibition hall, which '
        'opened for the first time at this year\'s conference. Over 85 partner companies and '
        'vendors had set up booths showcasing integrations, add-ons, and complementary solutions. '
        'The hall was organized into themed zones: Analytics & Reporting, Security & Compliance, '
        'Industry Solutions, and the popular Startup Showcase featuring 20 emerging companies.'
    )

    # ---- Paragraph 6: Expo Hall activities ----
    p6 = doc.add_paragraph(
        'Attendees could participate in live product demonstrations, enter raffles for technology '
        'prizes, and pick up continuing education credits at select booths. The partner pavilion '
        'proved to be one of the most visited sections of the conference, with booth staff '
        'reporting an average of 400-500 conversations per day. Several new partnership '
        'announcements were made on the expo floor, generating significant buzz on social media.'
    )

    # ---- Paragraph 7: Expo Hall photo marker ----
    p7 = doc.add_paragraph(
        'The exhibit hall bustled with energy from opening to close each day, with many attendees '
        'returning multiple times to revisit booths and complete in-depth product evaluations.'
    )

    # Photo marker 2 — plain text placeholder (to be replaced by agent)
    photo2 = doc.add_paragraph('[Photo: Expo Hall]')
    photo2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Paragraph 8: Technical sessions ----
    doc.add_heading('Technical Sessions and Workshops', level=2)
    p8 = doc.add_paragraph(
        'The conference featured five concurrent tracks running from 10:00 AM to 5:30 PM on Days '
        'One and Two: Platform Administration, Advanced Analytics, Developer APIs, Industry '
        'Verticals, and New User Foundations. This year, 34% of sessions were rated "Advanced" '
        'by attendees in post-session surveys, reflecting the growing sophistication of the user '
        'community. Workshops were limited to 40 participants to ensure hands-on time with '
        'instructors, and all 96 workshop slots sold out within 48 hours of registration opening.'
    )

    # Page break before Day 3 section
    doc.add_page_break()

    # ---- Paragraph 9: Award Ceremony intro ----
    doc.add_heading('The Pinnacle Awards Ceremony', level=2)
    p9 = doc.add_paragraph(
        'The highlight of Day Three was the annual Pinnacle Awards Gala, held in the Grand '
        'Ballroom on the evening of June 14. Over 800 attendees dressed in business-formal attire '
        'gathered to celebrate the outstanding achievements of customers, partners, and internal '
        'teams. The ceremony was hosted by comedian and tech enthusiast Diane Rosario, whose '
        'sharp wit and insider knowledge of enterprise software kept the audience thoroughly '
        'entertained between award presentations.'
    )

    # ---- Paragraph 10: Award Ceremony photo marker ----
    p10 = doc.add_paragraph(
        'Award categories included Innovation of the Year, Best ROI Story, Community Champion, '
        'Partner Excellence, and the inaugural Sustainability in Tech award. This year saw a '
        'record 312 nominations across all categories, reviewed by a panel of 24 independent '
        'judges drawn from industry analysts and academic institutions.'
    )

    # Photo marker 3 — plain text placeholder (to be replaced by agent)
    photo3 = doc.add_paragraph('[Photo: Award Ceremony]')
    photo3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ---- Remaining content ----
    p11 = doc.add_paragraph(
        'Winners were recognized with custom crystal trophies and featured in the company\'s '
        'annual impact report. Several award recipients gave brief acceptance speeches sharing '
        'how they had transformed their organizations using the platform. The evening concluded '
        'with a networking reception and live jazz band that continued well past midnight.'
    )

    doc.add_heading('Looking Ahead to 2026', level=2)
    p12 = doc.add_paragraph(
        'Before closing ceremonies, conference director Marcus Okafor announced that the 2026 '
        'Annual User Conference will be held in San Francisco, California, from May 18-20. Early '
        'bird registration will open November 1, 2025, with a 20% discount for current attendees '
        'who register before December 31. The announcement was met with enthusiastic applause, '
        'and the conference hashtag #UserConf2025 continued trending on social platforms for '
        'three days after the event concluded.'
    )

    p13 = doc.add_paragraph(
        'Survey results collected from 2,187 attendees (64% response rate) showed an overall '
        'satisfaction score of 4.7 out of 5.0, up from 4.5 in 2024. The most commonly cited '
        'strengths were session quality, networking opportunities, and venue logistics. Areas '
        'identified for improvement included Wi-Fi reliability in the overflow rooms and the '
        'need for more advanced-track content in the Analytics track. The conference team thanks '
        'all speakers, sponsors, and attendees for making the 2025 event our most successful to date.'
    )

    # Save to home directory
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also ensure Desktop directory exists and place file there as expected by task context
    desktop_dir = f'{WORKDIR}/Desktop'
    os.makedirs(desktop_dir, exist_ok=True)
    import shutil
    shutil.copy(OUTPUT, DESKTOP_PATH)
    print(f'File also copied to Desktop: {DESKTOP_PATH}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
