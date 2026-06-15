"""
Initial Setup: Anchor the image on page 1 to a specific character position
Task ID: writer_obj_045
Domain: libreoffice_writer

Creates a .docx file with an image currently anchored to (inline in) the first paragraph.
The agent must re-anchor the image to the beginning of the third paragraph with anchor type 'To Character'.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_045'
# Context says file is anchored_doc.docx at ~/Desktop/
DESKTOP = f'{WORKDIR}/Desktop'
FILENAME = 'anchored_doc.docx'
OUTPUT = f'{DESKTOP}/{FILENAME}'
# Also place at standard contract path
OUTPUT_STANDARD = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_image(path: str):
    """Create a simple placeholder image (5cm x 3cm equivalent in pixels)."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        # 5cm x 3cm at 96 DPI ~ 189x113 pixels
        img = Image.new('RGB', (189, 113), color=(70, 130, 180))
        draw = ImageDraw.Draw(img)
        draw.rectangle([5, 5, 183, 107], outline=(255, 255, 255), width=2)
        draw.text((45, 45), 'Company Logo', fill=(255, 255, 255))
        img.save(path)
    except ImportError:
        # Fallback: create minimal PNG
        import struct
        import zlib

        def png_chunk(chunk_type, data):
            c = chunk_type + data
            return struct.pack('>I', len(data)) + c + struct.pack('>I', zlib.crc32(c) & 0xFFFFFFFF)

        width, height = 189, 113
        ihdr = struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0)
        raw = b''
        for y in range(height):
            row = b'\x00'
            for x in range(width):
                row += bytes([70, 130, 180])
            raw += row
        compressed = zlib.compress(raw)
        png = (b'\x89PNG\r\n\x1a\n' +
               png_chunk(b'IHDR', ihdr) +
               png_chunk(b'IDAT', compressed) +
               png_chunk(b'IEND', b''))
        with open(path, 'wb') as f:
            f.write(png)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(DESKTOP, exist_ok=True)

    # Create temporary image file
    img_path = f'{WORKDIR}/company_logo_tmp.png'
    create_image(img_path)

    # Create document
    doc = Document()

    # --- Page title ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Nexus Technologies Inc.')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(12)

    # --- Paragraph 1: Image is anchored here (inline in this paragraph) ---
    para1 = doc.add_paragraph()
    # Add image inline at the start of paragraph 1 (simulating "anchored to paragraph 1")
    run_img = para1.add_run()
    run_img.add_picture(img_path, width=Cm(5), height=Cm(3))
    # Add text after image
    run_text1 = para1.add_run(
        ' Nexus Technologies is a leading provider of enterprise software solutions, '
        'serving clients across North America, Europe, and Asia-Pacific. Founded in 2008, '
        'the company has grown to over 2,500 employees worldwide.'
    )
    run_text1.font.size = Pt(11)
    para1.paragraph_format.space_after = Pt(8)

    # --- Paragraph 2 ---
    para2 = doc.add_paragraph(
        'Our mission is to empower businesses with cutting-edge technology that drives '
        'operational efficiency, enhances customer experiences, and accelerates digital '
        'transformation. We offer a comprehensive portfolio of products including ERP systems, '
        'CRM platforms, and cloud infrastructure services.'
    )
    para2.runs[0].font.size = Pt(11)
    para2.paragraph_format.space_after = Pt(8)

    # --- Paragraph 3: Agent should re-anchor image here (To Character) ---
    para3 = doc.add_paragraph(
        'The Nexus Innovation Lab, established in 2019, focuses on research and development '
        'of next-generation artificial intelligence and machine learning solutions. Our team '
        'of over 300 researchers collaborates with leading universities to push the boundaries '
        'of what\'s possible in enterprise AI.'
    )
    para3.runs[0].font.size = Pt(11)
    para3.paragraph_format.space_after = Pt(8)

    # --- Paragraph 4 ---
    para4 = doc.add_paragraph(
        'Nexus Technologies is committed to sustainability and corporate responsibility. '
        'In 2024, we achieved carbon neutrality across all our global operations, and we '
        'are on track to become net-zero by 2030. Our ESG initiatives have been recognized '
        'by several international organizations.'
    )
    para4.runs[0].font.size = Pt(11)
    para4.paragraph_format.space_after = Pt(8)

    # --- Paragraph 5 ---
    para5 = doc.add_paragraph(
        'With a strong financial track record—recording $1.2 billion in revenue for fiscal '
        'year 2024—Nexus Technologies continues to invest heavily in product development, '
        'talent acquisition, and strategic partnerships to maintain its competitive edge in '
        'the rapidly evolving technology landscape.'
    )
    para5.runs[0].font.size = Pt(11)

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also save to standard contract path
    import shutil
    shutil.copy(OUTPUT, OUTPUT_STANDARD)
    print(f'Also saved to: {OUTPUT_STANDARD}')

    # Clean up temp image
    if os.path.exists(img_path):
        os.remove(img_path)

    # GUI-ready startup: open file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
