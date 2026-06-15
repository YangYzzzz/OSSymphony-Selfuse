"""
Reward Script: Anchor image to character at beginning of third paragraph
Task ID: writer_obj_045
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): Image/drawing is in paragraph index 2 (the third paragraph)
  Component 2 (0.5): Drawing uses wp:anchor with positionH relativeFrom='character'
"""

import os

# Domain-specific imports
from docx import Document

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_obj_045'

# Namespace constant for wordprocessing drawing
WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def find_drawing_location(doc):
    """
    Find the paragraph index of the first drawing element in the document.
    Returns (para_index, run_index) or (None, None) if not found.
    """
    for i, para in enumerate(doc.paragraphs):
        for j, run in enumerate(para.runs):
            drawings = run._element.findall('.//{%s}drawing' % W_NS)
            if drawings:
                return i, j
    return None, None


def get_drawing_anchor_type(doc, para_index):
    """
    For a drawing at the given paragraph index, determine if it uses
    wp:anchor (floating) or wp:inline (inline).
    Returns: ('anchor', relativeFrom) or ('inline', None) or (None, None)
    """
    if para_index is None or para_index >= len(doc.paragraphs):
        return None, None

    para = doc.paragraphs[para_index]
    for run in para.runs:
        drawings = run._element.findall('.//{%s}drawing' % W_NS)
        for drawing in drawings:
            anchors = drawing.findall('{%s}anchor' % WP_NS)
            if anchors:
                for anchor in anchors:
                    posH = anchor.findall('{%s}positionH' % WP_NS)
                    for ph in posH:
                        relative_from = ph.get('relativeFrom')
                        return 'anchor', relative_from
                return 'anchor', None
            inlines = drawing.findall('{%s}inline' % WP_NS)
            if inlines:
                return 'inline', None
    return None, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.

    Task: Anchor the image on page 1 to a specific character position
          at the beginning of the third paragraph.
    - The image should be re-anchored to paragraph index 2 (third paragraph)
    - The anchor type should be 'To Character' (wp:anchor, relativeFrom='character')

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Image is in paragraph index 2 (the third paragraph) (0.5 points)
    # Initial env has the drawing at paragraph index 1 (second paragraph).
    # After the task, it should be at paragraph index 2 (third paragraph).
    try:
        para_idx, run_idx = find_drawing_location(doc)

        if para_idx is None:
            print("FAIL: Component 1 — No drawing/image found in the document")
        elif para_idx == 2:
            para_text = doc.paragraphs[para_idx].text[:60] if doc.paragraphs[para_idx].text else "(empty)"
            print(f"PASS: Component 1 — Drawing found at paragraph index 2 (third paragraph): '{para_text}...' (0.5 pts)")
            total_score += 0.5
        else:
            para_text = doc.paragraphs[para_idx].text[:60] if doc.paragraphs[para_idx].text else "(empty)"
            print(f"FAIL: Component 1 — Drawing at paragraph index {para_idx} (expected index 2). Para text: '{para_text}'")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Drawing uses wp:anchor with positionH relativeFrom='character' (0.5 points)
    # Initial env uses wp:inline. After the task, it should be wp:anchor anchored 'To Character'.
    # The key indicator of 'To Character' anchoring is positionH/@relativeFrom == 'character'.
    try:
        para_idx_for_anchor, _ = find_drawing_location(doc)
        anchor_type, relative_from = get_drawing_anchor_type(doc, para_idx_for_anchor)

        if anchor_type is None:
            print("FAIL: Component 2 — No drawing found to check anchor type")
        elif anchor_type == 'inline':
            print("FAIL: Component 2 — Drawing uses wp:inline, expected wp:anchor with relativeFrom='character'")
        elif anchor_type == 'anchor' and relative_from == 'character':
            print(f"PASS: Component 2 — Drawing uses wp:anchor with positionH relativeFrom='character' (0.5 pts)")
            total_score += 0.5
        elif anchor_type == 'anchor' and relative_from != 'character':
            print(f"FAIL: Component 2 — Drawing uses wp:anchor but positionH relativeFrom='{relative_from}', expected 'character'")
        else:
            print(f"FAIL: Component 2 — Unexpected anchor state: type={anchor_type}, relativeFrom={relative_from}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/anchored_doc.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
