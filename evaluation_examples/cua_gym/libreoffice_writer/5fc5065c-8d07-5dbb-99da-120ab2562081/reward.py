"""
Reward Script: Set up mirrored footer text in a Writer document
Task ID: writer_fs_083
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.25): evenAndOddHeaders is enabled in document settings
  - Component 2 (0.30): Odd (default/right) page footer contains 'Confidential' right-aligned
  - Component 3 (0.30): Even (left) page footer contains 'Confidential' left-aligned
  - Component 4 (0.15): Both footers have exactly 'Confidential' text (no extra text)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_083'


def persist_app_state(domain):
    """Save any unsaved changes in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as e:
        print(f"CRITICAL: Missing library: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    if len(doc.sections) == 0:
        print("CRITICAL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: evenAndOddHeaders is enabled (0.25 points)
    # This is a precondition that exists in initial too, BUT we need to confirm it's still
    # set after the agent acts. We combine it with a change-check below to avoid scoring
    # pre-existing state alone. Actually, since the initial also has this enabled, we should
    # NOT score this independently. Instead, we'll fold it into the footer content checks
    # as a gate condition.
    #
    # Revised: Use evenAndOddHeaders as a gate. If it's disabled, the even/odd footer
    # distinction doesn't work, so score 0.
    even_odd_enabled = 0
    try:
        settings_elem = doc.settings.element
        eaoh = settings_elem.findall('.//w:evenAndOddHeaders', nsmap)
        if len(eaoh) > 0:
            val = eaoh[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', None)
            # If val is None or not "false"/"0", it's enabled
            if val is None or val.lower() not in ('false', '0'):
                even_odd_enabled = 1
                print("GATE: evenAndOddHeaders is enabled - proceeding with checks")
            else:
                print("GATE FAIL: evenAndOddHeaders is disabled")
        else:
            print("GATE FAIL: evenAndOddHeaders element not found")
    except Exception as e:
        print(f"GATE ERROR: Could not check evenAndOddHeaders: {e}")

    if even_odd_enabled != 1:
        print("CRITICAL: Even/odd page footer distinction is not enabled")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Odd (default) page footer has 'Confidential' right-aligned (0.35 points)
    try:
        footer = section.footer
        odd_footer_text = ""
        odd_footer_alignment = None
        for p in footer.paragraphs:
            if p.text.strip():
                odd_footer_text = p.text.strip()
                odd_footer_alignment = p.paragraph_format.alignment
                break

        has_confidential_odd = odd_footer_text.lower() == 'confidential'
        is_right_aligned = odd_footer_alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT

        if has_confidential_odd and is_right_aligned:
            print(f"PASS: Component 1 - Odd page footer has 'Confidential' right-aligned (0.35 pts)")
            total_score += 0.35
        elif has_confidential_odd:
            print(f"FAIL: Component 1 - Odd page footer has 'Confidential' but alignment is {odd_footer_alignment}, expected RIGHT")
        elif odd_footer_text:
            print(f"FAIL: Component 1 - Odd page footer text is '{odd_footer_text}', expected 'Confidential'")
        else:
            print(f"FAIL: Component 1 - Odd page footer is empty")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Even page footer has 'Confidential' left-aligned (0.35 points)
    try:
        even_footer = section.even_page_footer
        even_footer_text = ""
        even_footer_alignment = None
        for p in even_footer.paragraphs:
            if p.text.strip():
                even_footer_text = p.text.strip()
                even_footer_alignment = p.paragraph_format.alignment
                break

        has_confidential_even = even_footer_text.lower() == 'confidential'
        # LEFT alignment can be represented as LEFT (0) or None (default is left)
        is_left_aligned = (even_footer_alignment == WD_PARAGRAPH_ALIGNMENT.LEFT or
                           even_footer_alignment is None)

        if has_confidential_even and is_left_aligned:
            print(f"PASS: Component 2 - Even page footer has 'Confidential' left-aligned (0.35 pts)")
            total_score += 0.35
        elif has_confidential_even:
            print(f"FAIL: Component 2 - Even page footer has 'Confidential' but alignment is {even_footer_alignment}, expected LEFT")
        elif even_footer_text:
            print(f"FAIL: Component 2 - Even page footer text is '{even_footer_text}', expected 'Confidential'")
        else:
            print(f"FAIL: Component 2 - Even page footer is empty")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Footer text mirrors correctly (opposite alignments) (0.30 points)
    # This checks the relationship: odd=RIGHT and even=LEFT simultaneously
    try:
        # Re-read both footers for the compound check
        odd_text = ""
        odd_align = None
        for p in section.footer.paragraphs:
            if p.text.strip():
                odd_text = p.text.strip()
                odd_align = p.paragraph_format.alignment
                break

        even_text = ""
        even_align = None
        for p in section.even_page_footer.paragraphs:
            if p.text.strip():
                even_text = p.text.strip()
                even_align = p.paragraph_format.alignment
                break

        odd_is_right = (odd_align == WD_PARAGRAPH_ALIGNMENT.RIGHT)
        even_is_left = (even_align == WD_PARAGRAPH_ALIGNMENT.LEFT or even_align is None)
        both_confidential = (odd_text.lower() == 'confidential' and
                             even_text.lower() == 'confidential')

        if both_confidential and odd_is_right and even_is_left:
            print(f"PASS: Component 3 - Footer text mirrors correctly on alternating pages (0.30 pts)")
            total_score += 0.30
        else:
            reasons = []
            if not both_confidential:
                reasons.append(f"text mismatch (odd='{odd_text}', even='{even_text}')")
            if not odd_is_right:
                reasons.append(f"odd not right-aligned (got {odd_align})")
            if not even_is_left:
                reasons.append(f"even not left-aligned (got {even_align})")
            print(f"FAIL: Component 3 - Mirror check failed: {', '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
