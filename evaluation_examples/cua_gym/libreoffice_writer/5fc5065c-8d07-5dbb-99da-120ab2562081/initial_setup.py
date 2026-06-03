"""
Initial Setup: Set up mirrored footers with 'Confidential' on alternating sides
Task ID: writer_fs_083
Domain: libreoffice_writer

Creates a 16-page document with mirrored page layout and empty footers.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_083'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup: mirrored margins (gutter) ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.gutter = Inches(0.5)

    # Enable mirrored margins via XML (mirrorMargins)
    sectPr = section._sectPr
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        # Setting gutter already done above; enable mirror margins at document level
        pass

    # Enable mirror margins in document settings
    settings = doc.settings.element
    mirror_el = settings.find(qn('w:mirrorMargins'))
    if mirror_el is None:
        mirror_el = settings.makeelement(qn('w:mirrorMargins'), {})
        settings.append(mirror_el)

    # Enable different odd/even headers/footers at document level
    even_odd = settings.find(qn('w:evenAndOddHeaders'))
    if even_odd is None:
        even_odd = settings.makeelement(qn('w:evenAndOddHeaders'), {})
        settings.append(even_odd)

    # --- Enable footers (empty) ---
    # Create empty odd page footer
    footer_odd = section.footer
    footer_odd.is_linked_to_previous = False
    if not footer_odd.paragraphs:
        footer_odd._element.append(footer_odd._element.makeelement(qn('w:p'), {}))
    # Clear any default text
    for p in footer_odd.paragraphs:
        p.text = ""

    # Access the even footer
    footer_even = section.even_page_footer
    footer_even.is_linked_to_previous = False
    if not footer_even.paragraphs:
        footer_even._element.append(footer_even._element.makeelement(qn('w:p'), {}))
    for p in footer_even.paragraphs:
        p.text = ""

    # --- Generate 16 pages of realistic content ---
    # Company quarterly report content
    chapters = [
        ("Q4 2025 Performance Summary", [
            "The fourth quarter of 2025 demonstrated strong performance across all business units, with consolidated revenue reaching $2.34 billion, representing a 12.7% year-over-year increase. Operating margins expanded by 180 basis points to 23.4%, driven primarily by operational efficiencies in our cloud services division.",
            "Our customer acquisition strategy yielded 14,200 new enterprise accounts during the quarter, bringing total active enterprise relationships to 87,500. Customer retention rates remained above 96%, reflecting the strength of our platform ecosystem and customer success programs.",
            "Capital expenditure for the quarter totaled $312 million, primarily allocated to data center expansion in the Asia-Pacific region and continued investment in our next-generation AI infrastructure platform.",
        ]),
        ("Revenue Analysis by Region", [
            "North America contributed $1.12 billion in revenue, up 9.3% from Q4 2024, driven by strong demand for our enterprise security solutions and expanded cloud migration services. The federal government segment showed particular strength, with contract wins totaling $178 million.",
            "Europe, Middle East, and Africa (EMEA) generated $687 million, reflecting 14.2% growth despite ongoing currency headwinds. The UK and German markets led performance, with combined revenue of $289 million. Our Paris data center, launched in September, is already operating at 67% capacity.",
            "Asia-Pacific revenue reached $389 million, a 22.1% increase driven by rapid adoption in Japan, South Korea, and emerging Southeast Asian markets. Our strategic partnership with NovaTech Solutions in Singapore has opened new distribution channels across the region.",
            "Latin America delivered $144 million in revenue, growing 18.5% year-over-year. Brazil and Mexico remained our largest markets, while Colombia and Chile showed promising early adoption patterns for our SMB-focused product suite.",
        ]),
        ("Product Division Highlights", [
            "The Cloud Infrastructure division posted $892 million in revenue, with our managed Kubernetes platform seeing 45% quarter-over-quarter growth in deployments. Enterprise adoption of our serverless computing framework accelerated, with over 3,200 production workloads migrated during the quarter.",
            "Our Enterprise Security suite generated $567 million, with the newly launched Zero Trust Network Access product contributing $89 million in its first full quarter. Threat detection response times improved by 34% following the integration of our proprietary ML models.",
            "Data Analytics and AI services contributed $481 million, up 31% year-over-year. Our AutoML platform now supports over 150 pre-built model templates, and customer-reported time-to-insight has decreased by an average of 62% compared to traditional approaches.",
            "The Collaboration and Productivity division reported $400 million in revenue. Our virtual workspace solution now supports 2.3 million daily active users across 12,000 organizations, with average session duration increasing to 6.2 hours per user per day.",
        ]),
        ("Financial Metrics and Projections", [
            "Gross profit for Q4 2025 was $1.59 billion, yielding a gross margin of 67.9%, up from 65.2% in the prior year period. This improvement reflects the increasing contribution of higher-margin SaaS revenue streams and continued optimization of our infrastructure costs.",
            "Research and development expenses totaled $423 million, representing 18.1% of revenue. Key investment areas included quantum-resistant encryption protocols, edge computing frameworks, and advanced natural language processing capabilities for our customer service automation platform.",
            "Free cash flow generation was $478 million for the quarter, bringing full-year 2025 free cash flow to $1.67 billion. The company repurchased $200 million in shares during the quarter and increased the quarterly dividend by 15% to $0.23 per share.",
            "For Q1 2026, management expects revenue in the range of $2.38 to $2.45 billion, representing 11-14% year-over-year growth. Operating margins are projected to remain stable at approximately 23-24%, with continued investment in AI and security research.",
        ]),
        ("Workforce and Operational Updates", [
            "Total headcount at quarter-end was 34,200, a net increase of 1,850 during the quarter. Engineering and product development accounted for 62% of new hires, with a particular focus on AI/ML specialists and cloud architecture roles.",
            "Our global office footprint expanded with the opening of innovation centers in Toronto and Bangalore, adding 45,000 square feet of collaborative workspace. The company's flexible work policy, allowing employees to work remotely up to three days per week, has contributed to a 23% reduction in voluntary attrition.",
            "Employee engagement scores, measured through quarterly pulse surveys, reached 84 out of 100, the highest level in company history. Key drivers included career development opportunities, competitive compensation adjustments made in Q3, and enhanced wellness benefits.",
        ]),
        ("Strategic Initiatives and Partnerships", [
            "The acquisition of CyberShield Analytics, completed in November 2025, has been fully integrated into our Enterprise Security division. The combined threat intelligence platform now processes over 4.2 billion security events daily, providing customers with industry-leading detection capabilities.",
            "Our partnership with GlobalTech Manufacturing to develop industrial IoT solutions entered its second phase, with pilot deployments at 15 manufacturing facilities across North America. Early results show 18% reduction in unplanned downtime and 12% improvement in energy efficiency.",
            "The company announced a strategic alliance with three major healthcare systems to develop HIPAA-compliant AI diagnostic support tools. This $150 million multi-year engagement represents our largest single healthcare sector commitment to date.",
            "Sustainability initiatives progressed significantly, with 78% of our data center operations now powered by renewable energy sources. We are on track to achieve our 100% renewable energy target by Q2 2027, ahead of the original 2028 deadline.",
        ]),
        ("Risk Factors and Compliance", [
            "The evolving regulatory landscape continues to present both challenges and opportunities. The EU AI Act implementation timeline requires modifications to certain automated decision-making features in our European product offerings, with compliance costs estimated at $45 million over the next 18 months.",
            "Cybersecurity threats remain elevated, with our security operations center detecting and mitigating 2.7 million attempted intrusions during the quarter. No material breaches occurred, and our incident response time averaged 4.2 minutes for critical severity events.",
            "Supply chain diversification efforts continued, reducing our dependence on any single semiconductor supplier to less than 30% of total chip procurement. New agreements with foundries in Taiwan, South Korea, and the United States provide improved resilience.",
            "Currency exposure management through natural hedging and forward contracts limited the negative foreign exchange impact to $23 million during the quarter, compared to $41 million in the prior year period.",
        ]),
        ("Outlook and Closing Remarks", [
            "As we enter 2026, the company is well-positioned to capitalize on secular growth trends in cloud computing, artificial intelligence, and cybersecurity. Our balanced portfolio, strong customer relationships, and continued investment in innovation provide a solid foundation for sustained growth.",
            "The Board of Directors has approved a new $1.5 billion share repurchase authorization, reflecting confidence in the company's long-term value creation potential. Combined with the increased dividend, total shareholder returns are expected to remain competitive.",
            "Management remains focused on disciplined capital allocation, operational excellence, and delivering innovative solutions that help our customers navigate an increasingly complex technology landscape. We are committed to creating long-term value for all stakeholders.",
            "We would like to thank our 34,200 employees worldwide for their dedication and contributions to another outstanding quarter. Their expertise, creativity, and commitment to customer success are the driving forces behind our continued growth and market leadership.",
        ]),
    ]

    for i, (title, paragraphs) in enumerate(chapters):
        # Add chapter heading
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.size = Pt(16)

        # Add content paragraphs
        for text in paragraphs:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"

        # Add page break after each chapter except the last
        if i < len(chapters) - 1:
            doc.add_page_break()

    # Add extra pages to reach 16 pages total (8 chapters with breaks = ~8+ pages)
    # Add appendices to fill remaining pages
    appendices = [
        ("Appendix A: Financial Statements Summary", [
            "The following tables present condensed consolidated financial statements for the quarter ended December 31, 2025. All figures are in millions of US dollars unless otherwise stated.",
            "Total Assets: $18,234M | Total Liabilities: $7,891M | Stockholders' Equity: $10,343M",
            "Current ratio improved to 2.1x from 1.9x in Q3, reflecting strong cash generation and disciplined working capital management. Debt-to-equity ratio remained stable at 0.34x.",
            "Deferred revenue increased by $234 million to $3.12 billion, indicating strong future revenue visibility. Remaining performance obligations totaled $8.7 billion as of December 31, 2025.",
        ]),
        ("Appendix B: Regional Office Directory", [
            "Corporate Headquarters: 1200 Innovation Boulevard, San Francisco, CA 94105, United States. Phone: +1 (415) 555-0100. This facility houses executive leadership, corporate strategy, and central engineering teams.",
            "European Headquarters: 45 Canary Wharf Tower, London E14 5AB, United Kingdom. Phone: +44 20 7555 0200. Regional operations span 14 countries with dedicated support centers in Frankfurt, Paris, and Amsterdam.",
            "Asia-Pacific Hub: 28 Marina Bay Financial Centre, Singapore 018982. Phone: +65 6555 0300. This office coordinates operations across 9 countries including Japan, South Korea, Australia, and India.",
            "Latin America Regional Office: Avenida Paulista 1578, Sao Paulo, SP 01310-200, Brazil. Phone: +55 11 5555 0400. This hub serves all Central and South American markets with multilingual support teams.",
        ]),
        ("Appendix C: Board of Directors", [
            "Dr. Elena Vasquez, Chair - Former CEO of TechVentures Global, with 30 years of experience in technology leadership. She holds a Ph.D. in Computer Science from MIT and has served on the boards of five Fortune 500 companies.",
            "Robert Tanaka, CEO - Appointed in 2021, previously served as COO for seven years. Under his leadership, the company has achieved 60% revenue growth and expanded into 12 new markets. MBA from Stanford Graduate School of Business.",
            "Priya Krishnamurthy, Independent Director - Managing Partner at Horizon Capital Partners. She brings deep expertise in financial strategy and capital markets, having led over $40 billion in technology sector transactions.",
            "James O'Sullivan, Independent Director - Retired Admiral, US Navy, and former Director of the National Cybersecurity Center. His defense and security background provides invaluable perspective on our enterprise security strategy.",
        ]),
        ("Appendix D: Glossary of Terms", [
            "ARR (Annual Recurring Revenue): The annualized value of active subscription contracts, providing a measure of predictable revenue streams. Our Q4 2025 ARR reached $8.2 billion.",
            "NRR (Net Revenue Retention): Measures revenue growth from existing customers including expansions and contractions. Our NRR of 118% indicates strong customer expansion and limited churn.",
            "DAU (Daily Active Users): The number of unique users who engage with our platform on a given day. Across all products, we averaged 4.7 million DAU in Q4 2025.",
            "MTTR (Mean Time to Resolution): The average time required to resolve customer support tickets. Our enterprise MTTR improved to 2.4 hours in Q4, down from 3.1 hours in Q3.",
        ]),
        ("Appendix E: Technology Roadmap Highlights", [
            "Our 2026 technology roadmap focuses on three strategic pillars: AI-native application development, zero-trust security architecture, and sustainable computing infrastructure.",
            "The next-generation platform, codenamed 'Horizon', will introduce breakthrough capabilities in real-time data processing, supporting up to 10 million events per second with sub-millisecond latency. Beta testing is scheduled for Q2 2026.",
            "Edge computing investments will extend our platform capabilities to 200+ edge locations globally by year-end 2026, enabling new use cases in autonomous systems, telemedicine, and smart infrastructure.",
            "Our quantum computing research lab, established in partnership with the University of Oxford, has achieved a significant milestone in error-corrected qubit operations, positioning us for practical quantum advantage in optimization workloads by 2028.",
        ]),
        ("Appendix F: Environmental, Social, and Governance Report", [
            "Carbon emissions decreased by 22% year-over-year, driven by renewable energy adoption and more efficient cooling systems in our data centers. Scope 1 and 2 emissions totaled 145,000 metric tons of CO2 equivalent.",
            "Water usage efficiency in data centers improved by 15% through implementation of advanced air cooling and water recycling systems. Our newest facilities in Oregon and Finland operate with a Power Usage Effectiveness (PUE) ratio of 1.08.",
            "Community investment programs reached $23 million in 2025, supporting STEM education, digital literacy initiatives, and disaster relief efforts across 45 countries. Over 8,000 employees volunteered a combined 120,000 hours.",
            "Board diversity metrics continued to improve, with women representing 40% of board members and underrepresented minorities comprising 30%. Our executive leadership team reflects similar diversity, with women holding 38% of VP-and-above positions.",
        ]),
        ("Appendix G: Patent and Intellectual Property Portfolio", [
            "The company was granted 342 new patents during Q4 2025, bringing our total active patent portfolio to 8,750. Key areas of patent activity included machine learning algorithms, distributed computing architectures, and biometric authentication methods.",
            "Licensing revenue from our intellectual property portfolio generated $67 million during the quarter. Cross-licensing agreements with major technology companies provide both revenue opportunities and freedom to operate in key technology domains.",
        ]),
        ("Appendix H: Contact Information", [
            "Investor Relations: Sarah Mitchell, VP Investor Relations. Email: ir@company.example.com. Phone: +1 (415) 555-0150. Quarterly earnings calls are typically held on the third Thursday following quarter-end.",
            "Media Inquiries: David Park, SVP Corporate Communications. Email: media@company.example.com. Phone: +1 (415) 555-0175.",
            "Customer Support: Available 24/7 via our online portal at support.company.example.com or by phone at +1 (800) 555-0199.",
        ]),
    ]

    for i, (title, paragraphs) in enumerate(appendices):
        doc.add_page_break()
        heading = doc.add_heading(title, level=2)
        for run in heading.runs:
            run.font.size = Pt(14)

        for text in paragraphs:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
