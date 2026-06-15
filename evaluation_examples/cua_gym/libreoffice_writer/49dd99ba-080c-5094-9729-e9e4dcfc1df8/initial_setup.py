"""
Initial Setup: Create a 16-page booklet document in A4 format for single-page printing.
Task ID: writer_rd_074
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_break(doc):
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # Set A4 page size with standard margins (single-page printing)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.orientation = WD_ORIENT.PORTRAIT

    # === PAGE 1: Cover Page ===
    doc.add_paragraph()  # spacer
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('Sustainable Urban Development', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Comprehensive Guide to Building Greener Cities')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.italic = True

    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('Prepared by the Metropolitan Planning Commission')
    run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('March 2025 Edition')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # === PAGE 2: Table of Contents ===
    add_page_break(doc)

    toc_heading = doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ('1. Introduction to Sustainable Urbanism', '3'),
        ('2. Green Infrastructure and Public Spaces', '4'),
        ('3. Renewable Energy Integration', '5'),
        ('4. Sustainable Transportation Networks', '6'),
        ('5. Water Management and Conservation', '7'),
        ('6. Waste Reduction and Circular Economy', '8'),
        ('7. Smart City Technologies', '9'),
        ('8. Community Engagement and Social Equity', '10'),
        ('9. Economic Considerations and Funding', '11'),
        ('10. Case Studies: Leading Green Cities', '12'),
        ('11. Implementation Roadmap', '13'),
        ('12. Monitoring and Evaluation Framework', '14'),
        ('13. Future Outlook and Recommendations', '15'),
    ]

    for title_text, page_num in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text} {"." * (50 - len(title_text))} {page_num}')
        run.font.size = Pt(11)

    # === PAGE 3: Chapter 1 - Introduction ===
    add_page_break(doc)
    doc.add_heading('1. Introduction to Sustainable Urbanism', level=1)
    doc.add_paragraph(
        'As the global population increasingly concentrates in urban areas, the need for '
        'sustainable development practices has never been more critical. By 2050, an estimated '
        '68% of the world\'s population will live in cities, placing unprecedented demands on '
        'infrastructure, resources, and ecosystems.'
    )
    doc.add_paragraph(
        'This guide presents a comprehensive framework for municipalities, urban planners, and '
        'community leaders seeking to transform their cities into models of sustainability. '
        'Drawing on best practices from cities worldwide, it covers key areas including green '
        'infrastructure, renewable energy, transportation, and community engagement.'
    )
    doc.add_paragraph(
        'The Metropolitan Planning Commission has developed this guide based on five years of '
        'research, field studies in over 40 cities, and consultation with leading experts in '
        'urban sustainability. Our goal is to provide actionable recommendations that can be '
        'adapted to diverse urban contexts.'
    )

    # === PAGE 4: Chapter 2 - Green Infrastructure ===
    add_page_break(doc)
    doc.add_heading('2. Green Infrastructure and Public Spaces', level=1)
    doc.add_paragraph(
        'Green infrastructure encompasses a network of natural and semi-natural areas designed '
        'to deliver ecosystem services. This includes urban forests, green roofs, rain gardens, '
        'bioswales, and permeable pavements that collectively manage stormwater, improve air '
        'quality, and enhance biodiversity.'
    )
    doc.add_paragraph(
        'Cities like Singapore, Copenhagen, and Portland have demonstrated that investing in '
        'green infrastructure yields significant returns. Singapore\'s "City in a Garden" '
        'initiative has increased green cover to 47% of the city\'s total land area, while '
        'Copenhagen\'s climate adaptation plan includes over 300 green infrastructure projects.'
    )
    doc.add_paragraph(
        'Key recommendations for municipalities include: establishing a minimum of 9 square '
        'meters of green space per capita, requiring green roofs on all new commercial buildings '
        'over 2,000 square meters, and creating connected greenway corridors linking major parks '
        'and natural areas.'
    )

    # === PAGE 5: Chapter 3 - Renewable Energy ===
    add_page_break(doc)
    doc.add_heading('3. Renewable Energy Integration', level=1)
    doc.add_paragraph(
        'Transitioning urban energy systems to renewable sources is fundamental to achieving '
        'carbon neutrality. Solar photovoltaics, wind power, geothermal systems, and district '
        'heating networks offer scalable solutions for urban energy demands.'
    )
    doc.add_paragraph(
        'Municipal governments can accelerate adoption through building codes requiring solar-ready '
        'design, community solar programs enabling renters to participate, and public-private '
        'partnerships for utility-scale installations. The cost of solar panels has decreased by '
        '89% since 2010, making rooftop solar increasingly accessible for residential properties.'
    )

    # === PAGE 6: Chapter 4 - Transportation ===
    add_page_break(doc)
    doc.add_heading('4. Sustainable Transportation Networks', level=1)
    doc.add_paragraph(
        'Transportation accounts for approximately 27% of urban greenhouse gas emissions globally. '
        'Shifting to sustainable mobility requires an integrated approach combining public transit '
        'expansion, cycling infrastructure, pedestrian-friendly design, and electric vehicle adoption.'
    )
    doc.add_paragraph(
        'Amsterdam\'s cycling infrastructure serves as a benchmark, with 767 kilometers of dedicated '
        'cycling paths carrying over 800,000 daily trips. Bogota\'s TransMilenio bus rapid transit '
        'system moves 2.4 million passengers daily at a fraction of the cost of subway construction.'
    )

    # === PAGE 7: Chapter 5 - Water Management ===
    add_page_break(doc)
    doc.add_heading('5. Water Management and Conservation', level=1)
    doc.add_paragraph(
        'Urban water management faces dual challenges: ensuring reliable supply while managing '
        'stormwater and wastewater. Climate change exacerbates both through altered precipitation '
        'patterns, rising temperatures, and more frequent extreme weather events.'
    )
    doc.add_paragraph(
        'Israel\'s national water strategy demonstrates what is achievable: through desalination, '
        'wastewater recycling (treating 86% for agricultural reuse), and demand management, the '
        'country has transformed from water scarcity to water security. Cities can implement similar '
        'strategies at the municipal level through greywater recycling mandates and smart metering.'
    )

    # === PAGE 8: Chapter 6 - Waste Reduction ===
    add_page_break(doc)
    doc.add_heading('6. Waste Reduction and Circular Economy', level=1)
    doc.add_paragraph(
        'The linear "take-make-dispose" model generates approximately 2.01 billion tonnes of '
        'municipal solid waste annually worldwide. Transitioning to a circular economy, where '
        'materials are continuously reused and recycled, is essential for urban sustainability.'
    )
    doc.add_paragraph(
        'San Francisco\'s Zero Waste program has achieved an 80% diversion rate through mandatory '
        'composting and recycling ordinances. Kamikatsu, Japan, separates waste into 45 categories '
        'and has achieved a 80% recycling rate. Extended producer responsibility legislation and '
        'pay-as-you-throw pricing create economic incentives for waste reduction.'
    )

    # === PAGE 9: Chapter 7 - Smart City Technologies ===
    add_page_break(doc)
    doc.add_heading('7. Smart City Technologies', level=1)
    doc.add_paragraph(
        'Digital technologies offer powerful tools for optimizing urban systems. Internet of Things '
        '(IoT) sensors, artificial intelligence, and data analytics enable real-time monitoring and '
        'management of energy grids, traffic flows, air quality, and water systems.'
    )
    doc.add_paragraph(
        'Barcelona\'s smart city platform integrates data from 19,500 sensors monitoring parking, '
        'lighting, waste management, and irrigation. The city estimates annual savings of EUR 75 '
        'million in water management alone. However, smart city initiatives must prioritize data '
        'privacy, cybersecurity, and equitable access to prevent digital divides.'
    )

    # === PAGE 10: Chapter 8 - Community Engagement ===
    add_page_break(doc)
    doc.add_heading('8. Community Engagement and Social Equity', level=1)
    doc.add_paragraph(
        'Sustainable urban development must center community participation and address systemic '
        'inequities. Environmental justice research consistently shows that low-income communities '
        'and communities of color bear disproportionate environmental burdens including air '
        'pollution, flooding risk, and heat island effects.'
    )
    doc.add_paragraph(
        'Participatory budgeting, as pioneered in Porto Alegre, Brazil, empowers residents to '
        'directly allocate portions of municipal budgets. Community land trusts protect affordable '
        'housing near transit corridors. Youth engagement programs, like those in Medellin, Colombia, '
        'build long-term civic capacity for sustainability leadership.'
    )

    # === PAGE 11: Chapter 9 - Economic Considerations ===
    add_page_break(doc)
    doc.add_heading('9. Economic Considerations and Funding', level=1)
    doc.add_paragraph(
        'Financing sustainable urban development requires creative approaches combining public '
        'funding, private investment, and innovative financial instruments. Green bonds have emerged '
        'as a major funding source, with global issuance exceeding USD 500 billion in 2023.'
    )
    doc.add_paragraph(
        'Tax increment financing (TIF) districts can capture property value increases generated by '
        'green infrastructure investments. Carbon pricing mechanisms, whether cap-and-trade or carbon '
        'taxes, generate revenue while incentivizing emissions reductions. The European Investment '
        'Bank\'s Smart Finance for Smart Buildings initiative demonstrates how public guarantees can '
        'leverage private capital for energy efficiency retrofits.'
    )

    # === PAGE 12: Chapter 10 - Case Studies ===
    add_page_break(doc)
    doc.add_heading('10. Case Studies: Leading Green Cities', level=1)

    doc.add_heading('Copenhagen, Denmark', level=2)
    doc.add_paragraph(
        'Copenhagen aims to become the world\'s first carbon-neutral capital by 2025. The city\'s '
        'strategy includes district heating from waste incineration, 450 km of cycling superhighways, '
        'and green roofs on all new buildings with roof pitches under 30 degrees.'
    )

    doc.add_heading('Curitiba, Brazil', level=2)
    doc.add_paragraph(
        'A pioneer in bus rapid transit since the 1970s, Curitiba\'s integrated transportation network '
        'carries 2.3 million passengers daily. The city\'s green exchange program allows low-income '
        'residents to trade recyclables for fresh produce and transit tokens.'
    )

    # === PAGE 13: Chapter 11 - Implementation Roadmap ===
    add_page_break(doc)
    doc.add_heading('11. Implementation Roadmap', level=1)
    doc.add_paragraph(
        'Successful implementation requires a phased approach. Phase 1 (Years 1-2) focuses on '
        'quick wins: LED streetlight conversion, cycling lane installation, and community garden '
        'programs. Phase 2 (Years 3-5) addresses systemic changes: building code updates, public '
        'transit expansion, and green infrastructure installation.'
    )
    doc.add_paragraph(
        'Phase 3 (Years 5-10) tackles transformative projects: district energy systems, complete '
        'streets redesign, and comprehensive stormwater management networks. Each phase should '
        'include measurable targets, allocated budgets, and designated responsible departments.'
    )

    # === PAGE 14: Chapter 12 - Monitoring Framework ===
    add_page_break(doc)
    doc.add_heading('12. Monitoring and Evaluation Framework', level=1)
    doc.add_paragraph(
        'A robust monitoring framework is essential for tracking progress and ensuring accountability. '
        'Key performance indicators should span environmental outcomes (CO2 emissions per capita, '
        'green space area), social equity (access to transit, affordable housing), and economic '
        'vitality (green jobs created, energy cost savings).'
    )
    doc.add_paragraph(
        'Annual sustainability reports should be published with standardized metrics aligned with '
        'the UN Sustainable Development Goals. Third-party audits every three years provide '
        'independent verification. Real-time dashboards powered by IoT sensors enable continuous '
        'monitoring of air quality, energy consumption, and water usage.'
    )

    # === PAGE 15: Chapter 13 - Future Outlook ===
    add_page_break(doc)
    doc.add_heading('13. Future Outlook and Recommendations', level=1)
    doc.add_paragraph(
        'The next decade will be decisive for urban sustainability. Emerging technologies including '
        'autonomous vehicles, advanced energy storage, vertical farming, and carbon capture present '
        'both opportunities and challenges for city planners.'
    )
    doc.add_paragraph(
        'Our key recommendations for municipal leaders are: establish legally binding carbon '
        'reduction targets, invest a minimum of 5% of municipal budgets in green infrastructure, '
        'create cross-departmental sustainability offices, and engage in regional and international '
        'knowledge-sharing networks.'
    )
    doc.add_paragraph(
        'The transition to sustainable cities is not merely an environmental imperative but an '
        'economic opportunity. Cities that lead in sustainability attract talent, investment, and '
        'innovation, creating a virtuous cycle of prosperity and environmental stewardship.'
    )

    # === PAGE 16: Back Cover ===
    add_page_break(doc)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    back_title = doc.add_paragraph()
    back_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = back_title.add_run('Metropolitan Planning Commission')
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()

    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = contact.add_run(
        '1200 Government Center Drive, Suite 450\n'
        'Metro City, MC 98001\n'
        'Tel: (555) 234-5678\n'
        'Email: sustainability@metroplanningcommission.gov\n'
        'Web: www.metroplanningcommission.gov/sustainability'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    doc.add_paragraph()

    copyright_para = doc.add_paragraph()
    copyright_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = copyright_para.add_run('\u00A9 2025 Metropolitan Planning Commission. All rights reserved.')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
