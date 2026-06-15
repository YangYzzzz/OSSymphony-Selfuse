"""
Initial Setup: Create a research document with hierarchical headings.
Task ID: writer_fp_029
Domain: libreoffice_writer

The document has 6 Heading 1 chapters. Chapter 3 (Methodology) contains
subsections including 'Data Analysis' at Heading 2 level with its own
Heading 3 sub-subsections.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_paragraph(doc, text):
    """Add a normal body paragraph with realistic content."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)
    return para


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Advanced Machine Learning Approaches for Environmental Monitoring', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Chapter 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    add_body_paragraph(doc,
        'Environmental monitoring has become increasingly critical in the face of accelerating '
        'climate change and biodiversity loss. This research explores the application of advanced '
        'machine learning techniques to improve the accuracy and efficiency of environmental '
        'monitoring systems across diverse ecosystems.')
    add_body_paragraph(doc,
        'The rapid development of sensor technologies and satellite imagery has generated '
        'unprecedented volumes of environmental data. Traditional analytical methods struggle '
        'to process this information at the required scale, creating an urgent need for '
        'automated approaches that can extract meaningful patterns from complex datasets.')

    # --- Chapter 2: Literature Review ---
    doc.add_heading('2. Literature Review', level=1)
    add_body_paragraph(doc,
        'Previous studies have demonstrated the potential of neural networks for species '
        'identification (Martinez et al., 2023) and habitat classification (Chen & Williams, 2024). '
        'However, significant gaps remain in real-time processing capabilities and cross-ecosystem '
        'generalization of trained models.')

    doc.add_heading('2.1 Remote Sensing Applications', level=2)
    add_body_paragraph(doc,
        'Remote sensing data from Landsat-8 and Sentinel-2 satellites has been widely used for '
        'land cover classification. Recent advances in convolutional neural networks have improved '
        'classification accuracy from 78% to 94% in tropical forest monitoring (Rivera et al., 2024).')

    doc.add_heading('2.2 Biodiversity Assessment Tools', level=2)
    add_body_paragraph(doc,
        'Automated species identification using acoustic monitoring and image recognition has '
        'shown promising results. The BirdNET framework achieved 89% accuracy across 984 species '
        'in field conditions (Kahl et al., 2023).')

    # --- Chapter 3: Methodology ---
    doc.add_heading('3. Methodology', level=1)
    add_body_paragraph(doc,
        'Our research methodology combines multiple data sources and analytical frameworks to '
        'create a comprehensive environmental monitoring pipeline. The approach integrates '
        'satellite imagery, ground-based sensors, and citizen science observations.')

    doc.add_heading('3.1 Data Collection', level=2)
    add_body_paragraph(doc,
        'Data was collected from 47 monitoring stations across three biomes: temperate forest, '
        'grassland, and wetland ecosystems. Each station recorded temperature, humidity, soil '
        'moisture, and acoustic data at 15-minute intervals from January 2024 to December 2025. '
        'A total of 2.3 million data points were gathered during the study period.')
    add_body_paragraph(doc,
        'Satellite imagery was acquired from the Copernicus Open Access Hub, providing '
        'multispectral data at 10-meter resolution for each monitoring site. Cloud-free composites '
        'were generated using median pixel selection over 30-day windows.')

    doc.add_heading('3.2 Data Analysis', level=2)
    add_body_paragraph(doc,
        'The analytical framework employed a multi-stage processing pipeline designed to handle '
        'heterogeneous data streams. Feature extraction, normalization, and model training were '
        'conducted using a distributed computing infrastructure with GPU acceleration.')

    doc.add_heading('3.2.1 Quantitative Methods', level=3)
    add_body_paragraph(doc,
        'Statistical analysis was performed using generalized linear mixed models (GLMMs) to '
        'account for spatial and temporal autocorrelation in the monitoring data. Machine learning '
        'models including Random Forest, XGBoost, and deep neural networks were trained on 80% '
        'of the dataset with 20% held out for validation. Cross-validation with 5 folds ensured '
        'robust performance estimates.')
    add_body_paragraph(doc,
        'Time series decomposition was applied to sensor data using seasonal-trend decomposition '
        'using LOESS (STL) to separate long-term trends from seasonal patterns and residual noise. '
        'Anomaly detection algorithms flagged observations exceeding 3 standard deviations.')

    doc.add_heading('3.2.2 Qualitative Methods', level=3)
    add_body_paragraph(doc,
        'Semi-structured interviews were conducted with 23 field researchers and conservation '
        'managers to understand current monitoring practices and identify gaps in existing tools. '
        'Thematic analysis of interview transcripts revealed five key themes: data accessibility, '
        'real-time alerting, integration challenges, training needs, and scalability concerns.')
    add_body_paragraph(doc,
        'Participatory mapping exercises with local communities provided valuable ground-truth '
        'data on land use changes and biodiversity observations that complemented the automated '
        'monitoring systems.')

    # --- Chapter 4: Results ---
    doc.add_heading('4. Results', level=1)
    add_body_paragraph(doc,
        'The integrated monitoring system demonstrated significant improvements over traditional '
        'methods. Overall species detection accuracy reached 91.3% across all three biomes, '
        'with particularly strong performance in wetland ecosystems (94.7%).')

    doc.add_heading('4.1 Model Performance', level=2)
    add_body_paragraph(doc,
        'The deep learning ensemble model outperformed individual classifiers with an F1 score '
        'of 0.923 on the held-out test set. Processing time for a single satellite image decreased '
        'from 47 minutes with traditional methods to 3.2 minutes with the optimized pipeline.')

    doc.add_heading('4.2 Ecosystem Health Indicators', level=2)
    add_body_paragraph(doc,
        'Novel composite indicators derived from the multi-modal data streams showed strong '
        'correlation (r = 0.87, p < 0.001) with established biodiversity indices. Early warning '
        'signals for ecosystem degradation were detected an average of 6.4 weeks before visible '
        'changes became apparent in satellite imagery.')

    # --- Chapter 5: Discussion ---
    doc.add_heading('5. Discussion', level=1)
    add_body_paragraph(doc,
        'The results confirm that machine learning approaches can substantially enhance '
        'environmental monitoring capabilities. The multi-modal data fusion strategy proved '
        'essential for achieving high accuracy, as no single data source provided sufficient '
        'information for reliable ecosystem assessment.')
    add_body_paragraph(doc,
        'Limitations of this study include the geographic restriction to temperate regions '
        'and the relatively short monitoring period. Future work should extend the framework '
        'to tropical and arctic ecosystems and incorporate longer temporal baselines to capture '
        'multi-year ecological dynamics.')

    # --- Chapter 6: Conclusion ---
    doc.add_heading('6. Conclusion', level=1)
    add_body_paragraph(doc,
        'This research demonstrates that integrated machine learning pipelines can transform '
        'environmental monitoring from reactive observation to proactive prediction. The open-source '
        'framework developed through this study is available for adoption by conservation agencies '
        'and research institutions worldwide.')
    add_body_paragraph(doc,
        'We recommend that environmental agencies invest in sensor infrastructure and computing '
        'capacity to support AI-driven monitoring systems. Training programs for field staff '
        'should be developed alongside technical deployments to ensure effective adoption and '
        'long-term sustainability of automated monitoring approaches.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
