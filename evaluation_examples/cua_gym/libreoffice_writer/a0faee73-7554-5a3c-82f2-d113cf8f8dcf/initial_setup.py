"""
Initial Setup: Employee handbook excerpt with 4 paragraphs all using single line spacing.
Task ID: osworld_writer_line_spacing_per_paragraph_006
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'osworld_writer_line_spacing_per_paragraph_006'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Paragraph 1: Introduction (single spacing) ---
    p1 = doc.add_paragraph(
        "Welcome to Meridian Solutions. This employee handbook outlines the policies, "
        "procedures, and expectations that guide our workplace. All employees are encouraged "
        "to read this document carefully and refer to it whenever questions arise about "
        "company guidelines. Our goal is to foster a collaborative, respectful, and "
        "productive environment for every member of our team."
    )
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(6)

    # --- Paragraph 2: Code of Conduct (single spacing - task will change to 1.5) ---
    p2 = doc.add_paragraph(
        "Our Code of Conduct establishes the standard of professional behavior expected "
        "from all Meridian Solutions employees. Employees must treat colleagues, clients, "
        "and partners with courtesy and respect at all times. Harassment, discrimination, "
        "or any form of misconduct will not be tolerated and may result in disciplinary "
        "action up to and including termination of employment."
    )
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)

    # --- Paragraph 3: Work Hours and Attendance (single spacing) ---
    p3 = doc.add_paragraph(
        "Standard working hours at Meridian Solutions are Monday through Friday, 9:00 AM "
        "to 5:30 PM. Employees are expected to maintain reliable attendance and to notify "
        "their manager at least one hour in advance if they are unable to report to work. "
        "Flexible work arrangements may be available with prior approval from your department "
        "head and Human Resources."
    )
    p3.paragraph_format.line_spacing = 1.0
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(6)

    # --- Paragraph 4: Benefits and Compensation (single spacing - task will change to 1.5) ---
    p4 = doc.add_paragraph(
        "Meridian Solutions offers a competitive benefits package to all full-time employees, "
        "including health, dental, and vision insurance, a 401(k) retirement plan with "
        "employer matching up to 4%, and 15 days of paid time off annually. Performance "
        "reviews are conducted semi-annually, and merit-based salary increases are considered "
        "at the completion of each review cycle."
    )
    p4.paragraph_format.line_spacing = 1.0
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
