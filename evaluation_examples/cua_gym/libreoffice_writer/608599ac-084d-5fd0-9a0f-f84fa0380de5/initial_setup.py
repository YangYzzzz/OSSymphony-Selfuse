"""
Initial Setup: Multi-level numbered list bylaws document
Task ID: wrpara_049
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'wrpara_049'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_level1(doc, number, text):
    """Add a Level 1 list item: '1. text'"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.0)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(f"{number}. {text}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Liberation Serif"
    return para


def add_level2(doc, letter, text):
    """Add a Level 2 list item: 'a. text'"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.4)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"{letter}. {text}")
    run.font.size = Pt(12)
    run.font.name = "Liberation Serif"
    return para


def add_level3(doc, numeral, text):
    """Add a Level 3 list item: 'i. text'"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.8)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"{numeral}. {text}")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"
    return para


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("Bylaws of the Greenfield Community Association", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Adopted March 15, 2025")
    run.font.size = Pt(11)
    run.font.name = "Liberation Serif"
    run.italic = True

    doc.add_paragraph()  # spacer

    # === Level 1, Item 1: Name and Purpose ===
    add_level1(doc, 1, "Name and Purpose")

    add_level2(doc, "a", "The name of this organization shall be the Greenfield Community Association, hereinafter referred to as the Association")
    add_level3(doc, "i", "The Association was originally incorporated under the laws of the State of Oregon on January 12, 2018")
    add_level3(doc, "ii", "All prior organizational documents are superseded by these Bylaws upon adoption")

    add_level2(doc, "b", "The purpose of the Association is to promote the welfare, safety, and beautification of the Greenfield residential community")
    add_level3(doc, "i", "This includes maintaining common areas, recreational facilities, and shared infrastructure")

    add_level2(doc, "c", "The Association shall operate as a nonprofit organization and shall not engage in activities for the private benefit of any individual member")

    # === Level 1, Item 2: Membership ===
    add_level1(doc, 2, "Membership")

    add_level2(doc, "a", "Membership in the Association is mandatory for all property owners within the Greenfield subdivision as recorded in the county assessor's office")
    add_level3(doc, "i", "Each residential lot shall be entitled to one membership regardless of the number of owners listed on the deed")
    add_level3(doc, "ii", "Commercial parcels located within the subdivision boundaries are also subject to mandatory membership")
    add_level3(doc, "iii", "Tenant occupants do not hold membership but may attend open meetings at the discretion of the Board")

    add_level2(doc, "b", "Annual membership dues shall be determined by the Board of Directors and approved by a majority vote at the annual general meeting")
    add_level3(doc, "i", "Dues must be paid in full by March 31 of each calendar year")
    add_level3(doc, "ii", "Late payments are subject to a penalty of 1.5 percent per month on the outstanding balance")

    add_level2(doc, "c", "A member may be suspended for nonpayment of dues exceeding ninety days, upon written notice delivered by certified mail")

    add_level2(doc, "d", "Reinstatement of a suspended member requires payment of all outstanding dues, penalties, and a reinstatement fee of fifty dollars")

    # === Level 1, Item 3: Board of Directors ===
    add_level1(doc, 3, "Board of Directors")

    add_level2(doc, "a", "The Board of Directors shall consist of seven members elected from the general membership at the annual meeting")
    add_level3(doc, "i", "Directors shall serve staggered terms of two years to ensure continuity of governance")
    add_level3(doc, "ii", "No director may serve more than three consecutive terms without a one-year hiatus")

    add_level2(doc, "b", "The Board shall elect from among its members a President, Vice President, Secretary, and Treasurer at the first meeting following the annual election")

    add_level2(doc, "c", "Vacancies on the Board may be filled by appointment of the remaining directors until the next annual election")
    add_level3(doc, "i", "An appointed director serves only the remainder of the vacated term")

    # === Level 1, Item 4: Meetings and Voting ===
    add_level1(doc, 4, "Meetings and Voting")

    add_level2(doc, "a", "The annual general meeting shall be held on the second Saturday of June each year at a location designated by the Board")
    add_level3(doc, "i", "Written notice of the annual meeting must be provided to all members at least thirty days in advance")
    add_level3(doc, "ii", "The notice must include the agenda, proposed budget, and any motions to be voted upon")

    add_level2(doc, "b", "Special meetings may be called by the President or upon written petition of at least twenty percent of the membership")

    add_level2(doc, "c", "A quorum for any general meeting shall consist of twenty-five percent of the total membership in good standing")
    add_level3(doc, "i", "If a quorum is not present, the meeting shall be adjourned and rescheduled within fourteen days")

    add_level2(doc, "d", "Each membership in good standing is entitled to one vote on all matters brought before the general membership")

    # === Level 1, Item 5: Amendments ===
    add_level1(doc, 5, "Amendments")

    add_level2(doc, "a", "These Bylaws may be amended by a two-thirds vote of the members present at any general or special meeting, provided a quorum is established")
    add_level3(doc, "i", "Proposed amendments must be submitted in writing to the Secretary at least forty-five days before the meeting at which they will be considered")
    add_level3(doc, "ii", "The Secretary shall distribute the proposed amendment text to all members at least thirty days before the vote")
    add_level3(doc, "iii", "Amendments affecting dues or assessments require an additional approval by simple majority of the full membership via written ballot")

    add_level2(doc, "b", "Amendments shall take effect immediately upon passage unless a different effective date is specified in the amendment itself")

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
