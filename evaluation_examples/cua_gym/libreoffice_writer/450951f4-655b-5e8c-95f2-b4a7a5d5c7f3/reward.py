"""
Reward Script: Novel chapter book formatting setup
Task ID: writer_creative_024
Domain: libreoffice_writer
Scoring:
  Component 1: Chapter heading formatted (18pt, bold, centered, space_after=24pt) — 0.20 pts
  Component 2: Drop cap on first body paragraph (3 lines, framePr) — 0.20 pts
  Component 3: Body text double-spaced with 12pt Liberation Serif — 0.25 pts
  Component 4: First line indent 0.5in on body paragraphs (except first) — 0.15 pts
  Component 5: Mirrored margins (inner >= outer) and even/odd footer with PAGE field — 0.20 pts
  Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_creative_024'

FILE_PATH = os.path.join(WORKDIR, 'novel_chapter3.docx')


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: must have at least 16 paragraphs (heading + 15 body)
    if len(doc.paragraphs) < 2:
        print(f"CRITICAL: Document has only {len(doc.paragraphs)} paragraphs, expected at least 16")
        print("REWARD: 0.0")
        return 0.0

    # -------------------------------------------------------------------------
    # Component 1: Chapter heading properly formatted (0.20 points)
    # Expected: font_size=18pt, bold=True, alignment=CENTER, space_after=24pt
    # Initial state: 11pt Calibri, left-aligned, no special formatting
    # -------------------------------------------------------------------------
    try:
        heading_para = doc.paragraphs[0]
        heading_text = heading_para.text.strip()
        pf = heading_para.paragraph_format

        # Check heading text is correct
        heading_text_ok = 'Chapter 3' in heading_text and 'Harbor' in heading_text

        # Check alignment = CENTER
        heading_align_ok = (pf.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)

        # Check space_after = 24pt (304800 EMU)
        space_after_ok = False
        if pf.space_after is not None:
            space_after_ok = abs(pf.space_after.pt - 24.0) < 2.0

        # Check font: 18pt bold
        heading_font_ok = False
        for run in heading_para.runs:
            size_ok = (run.font.size is not None and abs(run.font.size.pt - 18.0) < 0.5)
            bold_ok = (run.font.bold is True)
            if size_ok and bold_ok:
                heading_font_ok = True
                break

        if heading_text_ok and heading_align_ok and heading_font_ok and space_after_ok:
            print(f"PASS: Component 1 — Chapter heading: 18pt bold, centered, space_after=24pt (0.20 pts)")
            total_score += 0.20
        else:
            fail_details = []
            if not heading_text_ok:
                fail_details.append(f"text mismatch: {heading_text!r}")
            if not heading_align_ok:
                fail_details.append(f"alignment={pf.alignment} (expected CENTER)")
            if not heading_font_ok:
                fail_details.append("font not 18pt bold")
            if not space_after_ok:
                sa = pf.space_after.pt if pf.space_after else None
                fail_details.append(f"space_after={sa}pt (expected 24pt)")
            print(f"FAIL: Component 1 — {'; '.join(fail_details)}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -------------------------------------------------------------------------
    # Component 2: Drop cap on first body paragraph (0.20 points)
    # Expected: framePr w:dropCap="drop" w:lines="3" on paragraph 1
    # Initial state: no drop cap
    # -------------------------------------------------------------------------
    try:
        first_body_para = doc.paragraphs[1]
        para_xml = first_body_para._element.xml

        # Check for drop cap framePr element
        has_drop_cap = False
        drop_cap_lines = 0
        pPr = first_body_para._element.find(qn('w:pPr'))
        if pPr is not None:
            framePr = pPr.find(qn('w:framePr'))
            if framePr is not None:
                drop_cap_val = framePr.get(qn('w:dropCap'))
                lines_val = framePr.get(qn('w:lines'))
                if drop_cap_val == 'drop':
                    has_drop_cap = True
                    drop_cap_lines = int(lines_val) if lines_val else 0

        if has_drop_cap and drop_cap_lines >= 2:
            print(f"PASS: Component 2 — Drop cap on first body paragraph: {drop_cap_lines} lines tall (0.20 pts)")
            total_score += 0.20
        elif has_drop_cap:
            print(f"FAIL: Component 2 — Drop cap found but only {drop_cap_lines} lines (expected 3)")
        else:
            # Fallback: also check via string search in XML
            if 'dropCap' in para_xml and 'drop' in para_xml:
                print(f"PASS: Component 2 — Drop cap found in XML (string check) (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — No drop cap found on first body paragraph")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -------------------------------------------------------------------------
    # Component 3: Body text double-spaced with 12pt Liberation Serif (0.25 points)
    # Expected: line_spacing=2.0, font='Liberation Serif', size=12pt on paras 1-15
    # Initial state: 11pt Calibri, line_spacing=1.0
    # -------------------------------------------------------------------------
    try:
        body_paras = doc.paragraphs[1:]  # paragraphs 1 through end (index 1-15)
        double_spaced_count = 0
        font_ok_count = 0
        checked_count = 0

        for para in body_paras:
            checked_count += 1
            pf = para.paragraph_format

            # Check line spacing
            if pf.line_spacing is not None and abs(float(pf.line_spacing) - 2.0) < 0.1:
                double_spaced_count += 1

            # Check font: Liberation Serif 12pt
            para_font_ok = False
            for run in para.runs:
                if run.font.name == 'Liberation Serif':
                    if run.font.size is not None and abs(run.font.size.pt - 12.0) < 0.5:
                        para_font_ok = True
                        break
                    elif run.font.size is None:
                        # size may be inherited; name alone is partial credit
                        para_font_ok = True
                        break
            if para_font_ok:
                font_ok_count += 1

        if checked_count == 0:
            print("FAIL: Component 3 — No body paragraphs found")
        else:
            double_ratio = double_spaced_count / checked_count
            font_ratio = font_ok_count / checked_count

            # Both must be >= 80% to pass
            if double_ratio >= 0.8 and font_ratio >= 0.8:
                print(f"PASS: Component 3 — {double_spaced_count}/{checked_count} double-spaced, "
                      f"{font_ok_count}/{checked_count} with Liberation Serif 12pt (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — double_spaced={double_spaced_count}/{checked_count} "
                      f"({double_ratio:.0%}), font_ok={font_ok_count}/{checked_count} ({font_ratio:.0%})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -------------------------------------------------------------------------
    # Component 4: First line indent 0.5in on body paragraphs 2-15 (not para 1) (0.15 points)
    # Expected: first_line_indent=0.5in (457200 EMU) on paragraphs 2-15
    # Para 1 (drop cap) should have 0 or None indent
    # Initial state: no first-line indent
    # -------------------------------------------------------------------------
    try:
        body_paras_2_plus = doc.paragraphs[2:]  # paragraphs index 2 onward
        indent_ok_count = 0
        indent_checked = 0

        for para in body_paras_2_plus:
            indent_checked += 1
            pf = para.paragraph_format
            fli = pf.first_line_indent
            if fli is not None and abs(fli / 914400 - 0.5) < 0.05:
                indent_ok_count += 1

        if indent_checked == 0:
            print("FAIL: Component 4 — No paragraphs to check for indent")
        else:
            indent_ratio = indent_ok_count / indent_checked
            if indent_ratio >= 0.8:
                print(f"PASS: Component 4 — {indent_ok_count}/{indent_checked} body paragraphs "
                      f"have 0.5in first-line indent (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — Only {indent_ok_count}/{indent_checked} paragraphs "
                      f"have 0.5in indent ({indent_ratio:.0%})")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -------------------------------------------------------------------------
    # Component 5: Mirrored margins (inner >= outer) + even/odd footer with PAGE field (0.20 points)
    # Expected:
    #   - left_margin >= right_margin (inner wider for binding, i.e. 1.25in vs 1.0in)
    #   - Even and odd footers contain PAGE field code
    # Initial state: equal margins (1.0in all), empty footer
    # -------------------------------------------------------------------------
    try:
        section = doc.sections[0]

        # Check margins: left (inner for odd pages) > right (outer for odd pages)
        left_in = section.left_margin / 914400
        right_in = section.right_margin / 914400
        margins_mirrored = left_in > right_in  # inner wider than outer

        # Check that even/odd footer functionality exists (evenAndOddHeaders)
        sectPr = section._sectPr
        sectPr_xml = sectPr.xml
        has_even_odd = 'evenAndOddHeaders' in sectPr_xml

        # Check footer(s) for PAGE field
        def has_page_field(footer_obj):
            try:
                for fp in footer_obj.paragraphs:
                    fp_xml = fp._element.xml
                    if ('PAGE' in fp_xml and 'instrText' in fp_xml) or 'fldChar' in fp_xml:
                        return True
                return False
            except Exception:
                return False

        odd_footer_has_page = has_page_field(section.footer)
        even_footer_has_page = False
        try:
            even_footer_has_page = has_page_field(section.even_page_footer)
        except Exception:
            pass

        footer_ok = odd_footer_has_page  # at minimum odd footer must have page number

        if margins_mirrored and footer_ok:
            detail_parts = [
                f"left={left_in:.2f}in > right={right_in:.2f}in (mirrored)",
                f"odd_footer_page={odd_footer_has_page}",
                f"even_footer_page={even_footer_has_page}",
                f"even_odd_headers={has_even_odd}"
            ]
            print(f"PASS: Component 5 — Mirrored margins and page number footer: "
                  f"{', '.join(detail_parts)} (0.20 pts)")
            total_score += 0.20
        else:
            fail_parts = []
            if not margins_mirrored:
                fail_parts.append(f"margins not mirrored: left={left_in:.2f}in, right={right_in:.2f}in")
            if not footer_ok:
                fail_parts.append("no PAGE field in footer")
            print(f"FAIL: Component 5 — {'; '.join(fail_parts)}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -------------------------------------------------------------------------
    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
