"""
Initial Setup: Novel Chapter 3 - The Harbor
Task ID: writer_creative_024
Domain: libreoffice_writer

Creates a novel chapter document with:
- Chapter heading: "Chapter 3: The Harbor" (plain, unformatted)
- 15 paragraphs of realistic novel text (~3000 words)
- 11pt Calibri font, single-spaced, equal 1-inch margins
- No page numbers, no drop caps, no special formatting
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_creative_024'
OUTPUT = f'{WORKDIR}/Desktop/novel_chapter3.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Set default section margins: equal 1-inch margins (no mirroring)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Chapter heading — plain, no special formatting
    heading_para = doc.add_paragraph()
    heading_run = heading_para.add_run("Chapter 3: The Harbor")
    heading_run.font.name = "Calibri"
    heading_run.font.size = Pt(11)
    heading_run.bold = False
    heading_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    heading_para.paragraph_format.space_after = Pt(6)

    # 15 paragraphs of novel text — single-spaced, Calibri 11pt, no indent
    novel_paragraphs = [
        "The fog rolled in from the harbor before dawn, settling low over the weathered docks like a grey shroud. Elena Vasquez pulled her coat tighter and stared out at the invisible water, listening to the rhythmic slap of waves against the pilings and the distant cry of a gull somewhere above the mist.",
        "She had come back to Millhaven after twelve years, drawn by her aunt's letter and the faint scent of unfinished business. The harbor smelled the same: brine and diesel and something older, something that clung to the wood and stone and refused to let go. The fishing boats creaked in their berths, rocking gently as if in restless sleep.",
        "Thomas Aldridge met her at the end of Pier Seven, exactly as they had arranged. He looked older, the lines around his eyes deeper than she remembered, but his hands were the same — broad and steady, scarred from years of working the nets. He did not smile when he saw her.",
        '"You came," he said, his voice low and careful. "I wasn\'t sure you would."',
        '"I always come when the harbor calls," Elena replied, which was not entirely true, but felt right. She studied his face, looking for the lie behind the greeting. "How long has the boat been missing?"',
        "Three days, he told her. The Marianne had gone out on a Tuesday morning with a full crew of four and had not returned. The coast guard had searched for thirty-six hours and found nothing — no debris, no distress signal, no oil slick spreading on the grey water. Four men, gone, as if the sea had simply swallowed them whole.",
        "Elena followed Thomas down the length of the pier to a small office at the end, little more than a shack with a good view of the water. Inside, a hand-drawn chart of the harbor and the surrounding coastline covered most of one wall, marked with penciled X's and circles that meant nothing to her. A coffee percolator gurgled on a metal shelf, and the smell of burnt coffee cut through the salt air.",
        '"The last transmission was at 0347," Thomas said, moving to the chart. He pointed to a spot roughly two miles offshore, past the breakwater. "Normal position for that time of the run. Captain Reeves checked in at three forty-seven, reported weather holding, said they were working a good drift. Then nothing."',
        "She leaned in close to study the chart. The pencil marks were dense here, layered over months or years of careful notation. She could read the story in them if she looked long enough — the productive spots, the dangerous shoals, the routes the boats preferred at different tides. Her father had taught her to read charts like this, before everything went wrong.",
        '"Did Reeves have enemies?" she asked, knowing the question was predictable and necessary anyway.',
        '"Everybody has enemies," Thomas said. "Reeves had fewer than most. He was a straight dealer." He poured two cups of coffee without asking if she wanted one and set them on the chipped wooden desk. "But there was something he wanted to show me. The week before he went out, he told me he\'d found something. Said it would change everything."',
        "Outside, the fog was beginning to thin as the sun climbed somewhere above the overcast. Elena could make out the shapes of the nearest boats now — the Lady Ann, the Cormorant, old Haverford's trawler with its rusted wheelhouse. Normal harbor sounds drifted through the salt air: a radio playing somewhere, the clink of metal on metal, someone's boots on the dock.",
        '"What did he find?" she asked.',
        '"That\'s what I need you to figure out." Thomas wrapped both hands around his coffee cup and looked at her steadily. He had always been a man who said exactly what he meant and nothing more, which she had once found infuriating and now found almost restful. "You\'re the only person I trust who knows this harbor well enough to find what Reeves found. And I think whatever it was, it\'s why the Marianne didn\'t come back."',
        "Elena looked out through the smeared window at the harbor and the sea beyond it, grey and patient and completely indifferent to the small dramas of human beings played out on its surface. She thought about the twelve years she had spent away from this place, building a life on solid ground, telling herself she was done with mysteries and with water. She picked up her coffee cup. The harbor was calling, and she had always come when it called.",
    ]

    for para_text in novel_paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(para_text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        # Single-spacing, no indent
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open file in LibreOffice Writer with DISPLAY=:0
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
