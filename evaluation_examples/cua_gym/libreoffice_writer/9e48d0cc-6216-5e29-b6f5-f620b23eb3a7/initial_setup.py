"""
Initial Setup: Create a blank Writer master document (.odm) and three sub-documents
Task ID: writer_tech_053
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import zipfile

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_053'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_docx_with_content(filepath, title, paragraphs):
    """Create a .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)
        from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    for para_text in paragraphs:
        doc.add_paragraph(para_text)
    doc.save(filepath)
    print(f"Created: {filepath}")


def create_blank_odm(filepath):
    """Create a blank ODM (ODF Master Document) file."""
    manifest_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2">\n'
        ' <manifest:file-entry manifest:media-type="application/vnd.oasis.opendocument.text-master" manifest:version="1.2" manifest:full-path="/"/>\n'
        ' <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="content.xml"/>\n'
        ' <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="styles.xml"/>\n'
        '</manifest:manifest>'
    )

    content_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<office:document-content'
        ' xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
        ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"'
        ' xmlns:xlink="http://www.w3.org/1999/xlink"'
        ' office:version="1.2">\n'
        '  <office:body>\n'
        '    <office:text>\n'
        '      <text:p/>\n'
        '    </office:text>\n'
        '  </office:body>\n'
        '</office:document-content>'
    )

    styles_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<office:document-styles'
        ' xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"'
        ' xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"'
        ' xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"'
        ' office:version="1.2">\n'
        '  <office:styles>\n'
        '    <style:default-style style:family="paragraph">\n'
        '      <style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0.212cm"/>\n'
        '      <style:text-properties fo:font-size="12pt" style:font-name="Liberation Serif"/>\n'
        '    </style:default-style>\n'
        '  </office:styles>\n'
        '</office:document-styles>'
    )

    mimetype = 'application/vnd.oasis.opendocument.text-master'

    with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('mimetype', mimetype, compress_type=zipfile.ZIP_STORED)
        zf.writestr('META-INF/manifest.xml', manifest_xml)
        zf.writestr('content.xml', content_xml)
        zf.writestr('styles.xml', styles_xml)

    print(f"Created blank ODM: {filepath}")


def create_initial():
    # Install dependencies
    subprocess.run(['pip3', 'install', 'python-docx'], capture_output=True)

    # Create three sub-documents with realistic content
    create_docx_with_content(
        os.path.join(WORKDIR, 'chapter1_intro.docx'),
        'Chapter 1: Introduction',
        [
            'This chapter provides an overview of the project architecture and design principles.',
            'The system was designed with modularity and scalability in mind, following industry best practices for distributed computing environments.',
            'Key objectives include high availability, fault tolerance, and seamless horizontal scaling across multiple data centers.',
            'Throughout this document, we will explore the fundamental concepts that underpin the entire platform.',
        ]
    )

    create_docx_with_content(
        os.path.join(WORKDIR, 'chapter2_setup.docx'),
        'Chapter 2: Setup and Configuration',
        [
            'This chapter covers the installation and configuration procedures for the development environment.',
            'Prerequisites include Python 3.10 or later, Docker Engine 24.0+, and a minimum of 16GB RAM.',
            'Begin by cloning the repository and running the bootstrap script to initialize all required services.',
            'Configuration files are stored in the /etc/platform/ directory and use YAML format for readability.',
            'After completing the setup, run the validation suite to ensure all components are properly configured.',
        ]
    )

    create_docx_with_content(
        os.path.join(WORKDIR, 'chapter3_reference.docx'),
        'Chapter 3: API Reference',
        [
            'This chapter documents the complete API surface for the platform services.',
            'All endpoints follow RESTful conventions and return JSON responses with standard HTTP status codes.',
            'Authentication is handled via OAuth 2.0 bearer tokens with configurable expiration periods.',
            'Rate limiting is enforced at 1000 requests per minute per API key, with burst allowance of 50 requests.',
            'For detailed endpoint specifications, refer to the OpenAPI schema bundled with each service release.',
            'Error responses include a unique trace ID for debugging and correlation with server-side logs.',
        ]
    )

    # Create blank master document
    odm_path = os.path.join(WORKDIR, f'{TASK_ID}.odm')
    create_blank_odm(odm_path)

    # Launch LibreOffice Writer with the master document
    launch_gui(f'libreoffice --writer "{odm_path}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with blank master document')


create_initial()
