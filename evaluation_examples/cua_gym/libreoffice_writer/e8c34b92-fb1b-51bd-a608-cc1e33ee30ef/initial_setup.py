"""
Initial Setup: Legal document with TOC showing all 6 section headings
Task ID: writer_mt_069
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_069'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc(doc):
    """Insert a Table of Contents field that LibreOffice will update on open."""
    para = doc.add_paragraph()
    run = para.add_run()
    r_element = run._element

    # Begin field
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r_element.append(fldChar_begin)

    # Field instruction
    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    # Separate
    run3 = para.add_run()
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar_sep)

    # Placeholder TOC entries (will be updated by LibreOffice)
    toc_entries = ['Introduction', 'Facts', 'Legal Analysis', 'Arguments', 'Conclusion', 'Appendices']
    for entry in toc_entries:
        toc_para = doc.add_paragraph(entry)
        toc_para.style = doc.styles['Normal']
        toc_para.paragraph_format.space_after = Pt(2)

    # End field
    end_para = doc.add_paragraph()
    run_end = end_para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_end._element.append(fldChar_end)

    return para


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Legal Brief: Thompson v. Meridian Corp.', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Case No. 2025-CV-04892')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()  # spacing

    # --- Table of Contents ---
    toc_heading = doc.add_heading('Table of Contents', level=1)
    add_toc(doc)

    doc.add_page_break()

    # --- Section 1: Introduction ---
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'This legal brief is submitted on behalf of the Plaintiff, Emily Thompson, '
        'in the matter of Thompson v. Meridian Corporation. The Plaintiff seeks damages '
        'arising from breach of contract and negligent misrepresentation in connection '
        'with the commercial lease agreement executed on March 15, 2024.'
    )
    doc.add_paragraph(
        'The Plaintiff respectfully requests that this Court grant summary judgment in '
        'her favor based on the undisputed material facts and applicable law set forth '
        'herein. The evidence demonstrates that Meridian Corporation failed to fulfill '
        'its contractual obligations under Section 4.2 of the lease agreement.'
    )

    # --- Section 2: Facts ---
    doc.add_heading('Facts', level=1)
    doc.add_paragraph(
        'On March 15, 2024, the Plaintiff entered into a five-year commercial lease '
        'agreement with Meridian Corporation for Suite 1200 at 450 Commerce Boulevard, '
        'Springfield. The monthly rent was set at $8,750 with annual escalation of 3%. '
        'Meridian represented that the building met all current fire safety codes and '
        'had passed inspection within the preceding 12 months.'
    )
    doc.add_paragraph(
        'On July 22, 2024, the Springfield Fire Marshal conducted a routine inspection '
        'and found 14 code violations in the building, including non-functional sprinkler '
        'systems on floors 8 through 12 and blocked emergency exits in the parking garage. '
        'The Fire Marshal issued a Notice of Violation requiring remediation within 60 days.'
    )
    doc.add_paragraph(
        'Despite repeated written requests dated August 5, August 19, and September 3, '
        '2024, Meridian Corporation failed to commence any remediation work. On October 1, '
        '2024, the Plaintiff was forced to vacate the premises due to the unsafe conditions, '
        'incurring relocation costs of $34,500 and lost business revenue of $127,800.'
    )

    # --- Section 3: Legal Analysis ---
    doc.add_heading('Legal Analysis', level=1)
    doc.add_paragraph(
        'Under Springfield Revised Code Section 382.14, a commercial landlord has an '
        'affirmative duty to maintain leased premises in compliance with all applicable '
        'building and safety codes. Failure to do so constitutes a material breach of '
        'the lease agreement, entitling the tenant to terminate the lease and recover '
        'consequential damages.'
    )
    doc.add_paragraph(
        'The doctrine of negligent misrepresentation, as established in Harwood v. '
        'Pacific Realty (2019) and reaffirmed in Chen v. Metropolitan Properties (2022), '
        'holds that a landlord who makes affirmative representations about the condition '
        'of premises is liable for damages if those representations are made without '
        'reasonable basis. Here, Meridian represented that the building had passed fire '
        'inspection, when in fact the last inspection had occurred 26 months prior.'
    )

    # --- Section 4: Arguments ---
    doc.add_heading('Arguments', level=1)
    doc.add_paragraph(
        'First, the undisputed evidence establishes that Meridian Corporation breached '
        'the express terms of the lease agreement. Section 4.2 of the lease required '
        'Meridian to "maintain the premises in compliance with all applicable federal, '
        'state, and local laws, codes, and regulations." The 14 fire code violations '
        'documented by the Fire Marshal constitute a clear breach of this obligation.'
    )
    doc.add_paragraph(
        'Second, Meridian\'s pre-lease representations regarding fire safety compliance '
        'were materially false. The building had not been inspected within 12 months as '
        'represented; the last inspection occurred on January 8, 2022, more than two '
        'years before the lease was executed. This misrepresentation induced the Plaintiff '
        'to enter into the lease agreement and is actionable under both contract and tort law.'
    )
    doc.add_paragraph(
        'Third, the Plaintiff\'s damages are well-documented and directly traceable to '
        'Meridian\'s breach. The relocation costs of $34,500 and lost revenue of $127,800 '
        'are supported by receipts, invoices, and financial statements produced during '
        'discovery. The total damages sought amount to $162,300 plus pre-judgment interest.'
    )

    # --- Section 5: Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'For the foregoing reasons, the Plaintiff respectfully requests that this Court '
        'grant summary judgment in her favor and award damages in the amount of $162,300, '
        'together with pre-judgment interest at the statutory rate, costs of suit, and '
        'such other relief as this Court deems just and proper.'
    )
    doc.add_paragraph(
        'Respectfully submitted this 15th day of March, 2025.'
    )
    p_sig = doc.add_paragraph()
    p_sig.add_run('\n\n_______________________________\n').font.size = Pt(11)
    p_sig.add_run('Rachel M. Foster, Esq.\n').bold = True
    p_sig.add_run('Foster & Associates LLP\n')
    p_sig.add_run('Attorney for Plaintiff Emily Thompson')

    # --- Section 6: Appendices ---
    doc.add_heading('Appendices', level=1)
    doc.add_paragraph(
        'Appendix A: Commercial Lease Agreement dated March 15, 2024'
    )
    doc.add_paragraph(
        'Appendix B: Springfield Fire Marshal Notice of Violation dated July 22, 2024'
    )
    doc.add_paragraph(
        'Appendix C: Correspondence from Plaintiff to Meridian Corporation '
        '(August 5, August 19, September 3, 2024)'
    )
    doc.add_paragraph(
        'Appendix D: Relocation expense receipts and invoices totaling $34,500'
    )
    doc.add_paragraph(
        'Appendix E: Financial statements documenting lost revenue of $127,800'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=3.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
