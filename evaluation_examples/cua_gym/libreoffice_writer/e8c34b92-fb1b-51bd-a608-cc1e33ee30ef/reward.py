"""
Reward Script: Create a Table of Contents excluding the 'Appendices' heading
Task ID: writer_mt_069
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): 'Appendices' is NOT listed in the TOC region
  Component 2 (0.3): 'Appendices' heading outline level is overridden (excluded from TOC generation)
  Component 3 (0.3): The other 5 expected headings remain in the TOC
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_069'

# Expected TOC entries (all Heading 1 sections EXCEPT Appendices)
EXPECTED_TOC_ENTRIES = ['Introduction', 'Facts', 'Legal Analysis', 'Arguments', 'Conclusion']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # --- Locate TOC region ---
    # The TOC is between the "Table of Contents" Heading 1 and the next Heading
    toc_entries = []
    toc_started = False
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Heading 1' and para.text.strip() == 'Table of Contents':
            toc_started = True
            continue
        if toc_started:
            if para.style and para.style.name.startswith('Heading'):
                break
            text = para.text.strip()
            if text:
                toc_entries.append(text)

    print(f"INFO: Found TOC entries: {toc_entries}")

    # --- Component 1: 'Appendices' is NOT in the TOC (0.4 points) ---
    # This is the core task requirement. In initial_env, 'Appendices' IS in the TOC.
    # In golden_env, it should be removed.
    try:
        appendices_in_toc = any('Appendices' in entry for entry in toc_entries)
        if not appendices_in_toc and len(toc_entries) > 0:
            print(f"PASS: Component 1 - 'Appendices' is NOT in the TOC (0.4 pts)")
            total_score += 0.4
        elif appendices_in_toc:
            print(f"FAIL: Component 1 - 'Appendices' still appears in the TOC")
        else:
            print(f"FAIL: Component 1 - TOC region is empty (no entries found)")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # --- Component 2: 'Appendices' heading has outline level override to exclude from TOC (0.3 points) ---
    # The golden_env achieves exclusion by setting outlineLvl=9 on the Appendices heading.
    # This means the heading's outline level is set to body text level, excluding it from TOC.
    # Other valid approaches: changing style to non-heading, or outlineLvl to any value > typical TOC depth.
    try:
        appendices_heading = None
        for para in doc.paragraphs:
            if para.text.strip() == 'Appendices' and para.style and 'Heading' in para.style.name:
                appendices_heading = para
                break

        if appendices_heading is None:
            # Check if Appendices exists as body text (style changed away from Heading)
            for para in doc.paragraphs:
                if para.text.strip() == 'Appendices':
                    appendices_heading = para
                    break

        if appendices_heading is not None:
            pPr = appendices_heading._element.find('w:pPr', ns)
            # Check for outline level override
            outline_overridden = False
            if pPr is not None:
                outlineLvl = pPr.find('w:outlineLvl', ns)
                if outlineLvl is not None:
                    val = outlineLvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if val is not None and int(val) > 0:
                        outline_overridden = True
                        print(f"INFO: Appendices outlineLvl overridden to {val}")

            # Also check if style was changed to non-Heading
            style_changed = appendices_heading.style and 'Heading' not in appendices_heading.style.name

            if outline_overridden or style_changed:
                print(f"PASS: Component 2 - Appendices heading excluded from TOC via outline/style change (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 - Appendices heading still at default outline level with Heading style")
        else:
            print(f"FAIL: Component 2 - 'Appendices' paragraph not found in document body")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # --- Component 3: TOC has exactly 5 entries (the 5 expected headings, NOT Appendices) (0.3 points) ---
    # This component checks that the TOC has exactly the right count AND content.
    # In initial_env, there are 6 entries (including Appendices), so this FAILS on initial.
    # In golden_env, there are exactly 5 entries (excluding Appendices), so this PASSES on golden.
    try:
        found_count = 0
        missing = []
        for expected in EXPECTED_TOC_ENTRIES:
            if any(expected in entry for entry in toc_entries):
                found_count += 1
            else:
                missing.append(expected)

        # Must have exactly 5 entries (no more, no less) AND all 5 expected ones
        has_exactly_5 = len(toc_entries) == 5
        has_all_expected = found_count == len(EXPECTED_TOC_ENTRIES)

        if has_exactly_5 and has_all_expected:
            print(f"PASS: Component 3 - TOC has exactly 5 correct entries (0.3 pts)")
            total_score += 0.3
        elif has_all_expected and not has_exactly_5:
            print(f"FAIL: Component 3 - All 5 expected entries present but TOC has {len(toc_entries)} entries (expected 5)")
        else:
            print(f"FAIL: Component 3 - {found_count}/5 expected entries in TOC ({len(toc_entries)} total). Missing: {missing}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
