"""
Initial Setup: Insert a text frame in a newsletter document
Task ID: writer_fs_009
Domain: libreoffice_writer

Creates a newsletter document with title and body text. No text frame present.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_009'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Page setup - standard A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Title
    title = doc.add_heading("Greenfield Community Newsletter", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / edition info
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("Volume 12, Issue 4  |  April 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sep_run = sep.add_run("_" * 60)
    sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Section 1
    h1 = doc.add_heading("Community Garden Opens for Spring Planting", level=1)

    p1 = doc.add_paragraph()
    p1_run = p1.add_run(
        "The Greenfield Community Garden at Maple Park is now accepting registrations "
        "for the 2026 spring planting season. Residents can reserve individual plots "
        "measuring 3 meters by 4 meters for an annual fee of $45. This year, the garden "
        "committee has expanded the total area to accommodate 28 additional plots, "
        "bringing the total capacity to 150 garden spaces."
    )
    p1_run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        "Garden coordinator Elena Vasquez noted that last year saw a 35% increase in "
        "participation compared to 2024. \"We are thrilled by the growing interest in "
        "urban gardening,\" she said during the March town hall meeting. \"Our composting "
        "program alone diverted over 2.8 tons of organic waste from the landfill.\""
    )
    p2_run.font.size = Pt(11)

    # Section 2
    h2 = doc.add_heading("Library Renovation Update", level=1)

    p3 = doc.add_paragraph()
    p3_run = p3.add_run(
        "Phase two of the Greenfield Public Library renovation is on schedule for "
        "completion by July 15, 2026. The $3.2 million project includes a new children's "
        "reading wing, upgraded HVAC systems, and a digital media lab equipped with "
        "12 workstations for residents to access creative software and 3D printing services."
    )
    p3_run.font.size = Pt(11)

    p4 = doc.add_paragraph()
    p4_run = p4.add_run(
        "During construction, temporary library services continue to operate from the "
        "community center on Oak Street. Head librarian Thomas Park confirmed that the "
        "summer reading program will proceed as planned, with registration opening on "
        "May 1. Over 340 children participated in last year's program, logging a combined "
        "total of 12,500 reading hours."
    )
    p4_run.font.size = Pt(11)

    # Section 3
    h3 = doc.add_heading("Upcoming Events", level=1)

    events = [
        "April 12 - Spring Farmers Market Opening (Town Square, 8:00 AM - 1:00 PM)",
        "April 19 - Earth Day Cleanup Drive (meet at Riverside Park, 9:00 AM)",
        "April 26 - Community Potluck & Talent Show (Recreation Center, 5:30 PM)",
        "May 3 - Annual 5K Fun Run / Walk benefiting Greenfield Youth Programs",
        "May 10 - Mother's Day Craft Workshop at the Community Center (free, all ages)",
    ]
    for event in events:
        bp = doc.add_paragraph(event, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(11)

    # Closing
    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    closing_run = closing.add_run(
        "\nFor more information, visit www.greenfieldcommunity.org or call (555) 234-8901."
    )
    closing_run.font.size = Pt(10)
    closing_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    closing_run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
