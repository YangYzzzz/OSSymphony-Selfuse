"""
Initial Setup: Make the header row text bold and white, with a dark navy blue background. Also center-align all header cells horizontally.
Task ID: writer_tbl_041
Domain: libreoffice_writer

Creates: /home/user/Desktop/styled_report.docx
  - A table with 5 rows and 4 columns
  - Row 1 (header): 'Region' | 'Revenue' | 'Costs' | 'Profit'
  - Data rows 2-5: North/South/East/West regional financial data
  - All text is default (black, regular weight, left-aligned) — NO bold, NO background color
"""

import os
import shlex
import subprocess
import sys
import time

# Ensure python-docx is available on the VM
subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'], check=True)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_041'
OUTPUT = f'{WORKDIR}/styled_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default empty paragraph if it exists
    # Add the table directly
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # --- Header row (Row 1) ---
    header_data = ['Region', 'Revenue', 'Costs', 'Profit']
    for col_idx, header_text in enumerate(header_data):
        cell = table.cell(0, col_idx)
        # Clear default content and set plain text
        para = cell.paragraphs[0]
        # Left-aligned (default), no bold, no background
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(header_text)
        run.bold = False
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # black text

    # --- Data rows ---
    data_rows = [
        ['North', '$50,000', '$30,000', '$20,000'],
        ['South', '$45,000', '$28,000', '$17,000'],
        ['East',  '$60,000', '$35,000', '$25,000'],
        ['West',  '$55,000', '$32,000', '$23,000'],
    ]

    for row_idx, row_data in enumerate(data_rows, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = para.add_run(cell_text)
            run.bold = False

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
