"""
Reward Script: Make header row text bold and white, with dark navy blue background and center-aligned.
Task ID: writer_tbl_041
Domain: libreoffice_writer
Scoring:
  Component 1: Header row text is bold (all 4 cells)                     — 0.25 points
  Component 2: Header row text color is white (RGB: 255,255,255)         — 0.25 points
  Component 3: Header cell background is dark navy blue (~RGB: 0,0,128)  — 0.30 points
  Component 4: Header cells are center-aligned horizontally               — 0.20 points
  Total: 1.0
"""

import os
from math import sqrt

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_tbl_041'
FILE_NAME = 'styled_report.docx'


def color_distance(rgb_obj, target_tuple):
    """Compute Euclidean RGB distance between an RGBColor object and an (r,g,b) tuple."""
    return sqrt(
        (rgb_obj[0] - target_tuple[0]) ** 2 +
        (rgb_obj[1] - target_tuple[1]) ** 2 +
        (rgb_obj[2] - target_tuple[2]) ** 2
    )


def hex_to_rgb(hex_str):
    """Convert a 6-char hex string (e.g. '000080') to (r,g,b) tuple."""
    h = hex_str.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document — precondition gate
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Locate the first table — precondition gate
    try:
        if not doc.tables:
            print("CRITICAL: No tables found in document")
            print("REWARD: 0.0")
            return 0.0
        table = doc.tables[0]
        if len(table.rows) < 1:
            print("CRITICAL: Table has no rows")
            print("REWARD: 0.0")
            return 0.0
        header_row = table.rows[0]
        header_cells = header_row.cells
        num_header_cells = len(header_cells)
        print(f"INFO: Found table with {len(table.rows)} rows, {num_header_cells} header cells")
    except Exception as e:
        print(f"CRITICAL: Cannot access table: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Header row text is bold (0.25 points)
    # FAILS on initial (bold=False) → PASSES on golden (bold=True)
    try:
        bold_count = 0
        bold_details = []
        for c_idx, cell in enumerate(header_cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        is_bold = run.font.bold is True
                        bold_details.append(f"  cell[{c_idx}] run='{run.text}' bold={run.font.bold}")
                        if is_bold:
                            bold_count += 1

        if bold_count == num_header_cells:
            print(f"PASS: Component 1 — All {num_header_cells} header cells have bold text (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Only {bold_count}/{num_header_cells} header cells are bold")
            for d in bold_details:
                print(d)
    except Exception as e:
        print(f"ERROR: Component 1 (bold check) — {e}")

    # Component 2: Header row text color is white (RGB: 255,255,255) (0.25 points)
    # FAILS on initial (color=000000 or None) → PASSES on golden (color=FFFFFF)
    try:
        white_count = 0
        color_details = []
        WHITE_RGB = (255, 255, 255)
        COLOR_THRESHOLD = 30  # allow slight variation

        for c_idx, cell in enumerate(header_cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        try:
                            color_rgb = run.font.color.rgb
                            if color_rgb is not None:
                                dist = color_distance(color_rgb, WHITE_RGB)
                                color_details.append(
                                    f"  cell[{c_idx}] run='{run.text}' color=({color_rgb[0]},{color_rgb[1]},{color_rgb[2]}) dist={dist:.1f}"
                                )
                                if dist <= COLOR_THRESHOLD:
                                    white_count += 1
                            else:
                                color_details.append(f"  cell[{c_idx}] run='{run.text}' color=None (not set)")
                        except Exception as ce:
                            color_details.append(f"  cell[{c_idx}] run='{run.text}' color error: {ce}")

        if white_count == num_header_cells:
            print(f"PASS: Component 2 — All {num_header_cells} header cells have white text (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 — Only {white_count}/{num_header_cells} header cells have white text")
            for d in color_details:
                print(d)
    except Exception as e:
        print(f"ERROR: Component 2 (white color check) — {e}")

    # Component 3: Header cell background is dark navy blue (approximately RGB: 0,0,128) (0.30 points)
    # FAILS on initial (no background shading) → PASSES on golden (fill=000080)
    try:
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        NAVY_RGB = (0, 0, 128)
        BG_THRESHOLD = 60  # allow a wider range of navy blue shades

        navy_count = 0
        bg_details = []

        for c_idx, cell in enumerate(header_cells):
            shd = cell._tc.find(f'{{{W_NS}}}tcPr/{{{W_NS}}}shd')
            if shd is None:
                # Try direct child lookup
                shd = cell._tc.find(f'.//{{{W_NS}}}shd')

            if shd is not None:
                fill_hex = shd.get(f'{{{W_NS}}}fill')
                if fill_hex and fill_hex.upper() not in ('AUTO', 'FFFFFF', ''):
                    try:
                        fill_rgb = hex_to_rgb(fill_hex)
                        dist = color_distance(fill_rgb, NAVY_RGB)
                        bg_details.append(
                            f"  cell[{c_idx}] fill={fill_hex} rgb={fill_rgb} dist_from_navy={dist:.1f}"
                        )
                        if dist <= BG_THRESHOLD:
                            navy_count += 1
                        else:
                            bg_details.append(f"    -> dist {dist:.1f} exceeds threshold {BG_THRESHOLD}")
                    except Exception as he:
                        bg_details.append(f"  cell[{c_idx}] fill={fill_hex} parse error: {he}")
                else:
                    bg_details.append(f"  cell[{c_idx}] fill={fill_hex} (no meaningful background)")
            else:
                bg_details.append(f"  cell[{c_idx}] no shading element found")

        if navy_count == num_header_cells:
            print(f"PASS: Component 3 — All {num_header_cells} header cells have dark navy blue background (0.30 pts)")
            total_score += 0.30
        else:
            print(f"FAIL: Component 3 — Only {navy_count}/{num_header_cells} header cells have dark navy blue background")
            for d in bg_details:
                print(d)
    except Exception as e:
        print(f"ERROR: Component 3 (background color check) — {e}")

    # Component 4: Header cells are center-aligned horizontally (0.20 points)
    # FAILS on initial (LEFT alignment) → PASSES on golden (CENTER alignment)
    try:
        center_count = 0
        align_details = []

        for c_idx, cell in enumerate(header_cells):
            for para in cell.paragraphs:
                alignment = para.paragraph_format.alignment
                align_details.append(f"  cell[{c_idx}] alignment={alignment}")
                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    center_count += 1
                    break  # Only need one centered paragraph per cell

        if center_count == num_header_cells:
            print(f"PASS: Component 4 — All {num_header_cells} header cells are center-aligned (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 4 — Only {center_count}/{num_header_cells} header cells are center-aligned")
            for d in align_details:
                print(d)
    except Exception as e:
        print(f"ERROR: Component 4 (alignment check) — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: test against canonical artifact path
file_path = os.path.join(WORKDIR, FILE_NAME)
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
