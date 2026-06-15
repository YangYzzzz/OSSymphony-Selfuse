"""
Initial Setup: NDA Template with merge fields and data source
Task ID: writer_mt_046
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
import csv

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_046'
TEMPLATE_FILE = f'{WORKDIR}/NDA_Template.docx'
DATA_FILE = f'{WORKDIR}/NDAParties.csv'
OUTPUT_DIR = f'{WORKDIR}/Desktop/NDA_Output'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# NDA parties data - 8 records
NDA_PARTIES = [
    {"PartyName": "Elena Vasquez", "CompanyName": "Meridian Dynamics LLC", "EffectiveDate": "January 15, 2026", "Jurisdiction": "State of California"},
    {"PartyName": "James Whitfield", "CompanyName": "Apex Innovations Inc.", "EffectiveDate": "February 3, 2026", "Jurisdiction": "State of New York"},
    {"PartyName": "Priya Sharma", "CompanyName": "NovaTech Solutions Pvt. Ltd.", "EffectiveDate": "March 10, 2026", "Jurisdiction": "State of Delaware"},
    {"PartyName": "Marcus Chen", "CompanyName": "Silverline Partners Group", "EffectiveDate": "April 1, 2026", "Jurisdiction": "State of Texas"},
    {"PartyName": "Olivia Brennan", "CompanyName": "Coastal Ventures Corp.", "EffectiveDate": "May 22, 2026", "Jurisdiction": "State of Florida"},
    {"PartyName": "David Kowalski", "CompanyName": "Pinnacle Research Labs", "EffectiveDate": "June 8, 2026", "Jurisdiction": "State of Illinois"},
    {"PartyName": "Aisha Patel", "CompanyName": "Horizon Analytics Ltd.", "EffectiveDate": "July 14, 2026", "Jurisdiction": "State of Massachusetts"},
    {"PartyName": "Robert Lindqvist", "CompanyName": "Nordic Bridge Consulting AB", "EffectiveDate": "August 30, 2026", "Jurisdiction": "State of Washington"},
]


def create_nda_template():
    """Create the NDA template document with merge fields."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    title = doc.add_heading('NON-DISCLOSURE AGREEMENT', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Preamble
    preamble = doc.add_paragraph()
    preamble.paragraph_format.space_before = Pt(12)
    preamble.paragraph_format.space_after = Pt(6)
    run = preamble.add_run(
        'This Non-Disclosure Agreement ("Agreement") is entered into as of '
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    field_run = preamble.add_run('<EffectiveDate>')
    field_run.font.size = Pt(11)
    field_run.font.name = 'Times New Roman'
    field_run.bold = True

    run2 = preamble.add_run(' (the "Effective Date"), by and between:')
    run2.font.size = Pt(11)
    run2.font.name = 'Times New Roman'

    # Party info
    party_para = doc.add_paragraph()
    party_para.paragraph_format.space_before = Pt(6)
    party_para.paragraph_format.space_after = Pt(6)
    party_para.paragraph_format.left_indent = Inches(0.5)

    r = party_para.add_run('Party: ')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    r.bold = True

    r2 = party_para.add_run('<PartyName>')
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'
    r2.bold = True

    r3 = party_para.add_run(', representing ')
    r3.font.size = Pt(11)
    r3.font.name = 'Times New Roman'

    r4 = party_para.add_run('<CompanyName>')
    r4.font.size = Pt(11)
    r4.font.name = 'Times New Roman'
    r4.bold = True

    r5 = party_para.add_run(' (the "Receiving Party")')
    r5.font.size = Pt(11)
    r5.font.name = 'Times New Roman'

    # And
    and_para = doc.add_paragraph()
    and_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    and_para.paragraph_format.space_before = Pt(6)
    and_para.paragraph_format.space_after = Pt(6)
    r = and_para.add_run('AND')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    r.bold = True

    # Disclosing party
    disc_para = doc.add_paragraph()
    disc_para.paragraph_format.space_before = Pt(6)
    disc_para.paragraph_format.space_after = Pt(12)
    disc_para.paragraph_format.left_indent = Inches(0.5)
    r = disc_para.add_run('Quantum Cipher Technologies, Inc.')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    r.bold = True
    r2 = disc_para.add_run(' (the "Disclosing Party")')
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'

    # Section 1
    doc.add_heading('1. Definition of Confidential Information', level=2)
    s1 = doc.add_paragraph()
    r = s1.add_run(
        '"Confidential Information" means any and all non-public information, '
        'including but not limited to trade secrets, proprietary data, business plans, '
        'financial records, technical specifications, customer lists, marketing strategies, '
        'and any other information disclosed by the Disclosing Party to the Receiving Party, '
        'whether in writing, orally, electronically, or by any other means.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    # Section 2
    doc.add_heading('2. Obligations of the Receiving Party', level=2)
    s2 = doc.add_paragraph()
    r = s2.add_run(
        'The Receiving Party agrees to: (a) hold the Confidential Information in strict confidence; '
        '(b) not disclose the Confidential Information to any third party without the prior written '
        'consent of the Disclosing Party; (c) use the Confidential Information solely for the purpose '
        'of evaluating or engaging in business discussions with the Disclosing Party; and '
        '(d) take all reasonable measures to protect the secrecy of the Confidential Information.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    # Section 3
    doc.add_heading('3. Term and Termination', level=2)
    s3 = doc.add_paragraph()
    r = s3.add_run(
        'This Agreement shall remain in effect for a period of three (3) years from the Effective Date, '
        'unless earlier terminated by either party upon thirty (30) days written notice to the other party. '
        'The obligations of confidentiality shall survive the termination of this Agreement for a period '
        'of five (5) years.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    # Section 4
    doc.add_heading('4. Governing Law', level=2)
    s4 = doc.add_paragraph()
    r = s4.add_run(
        'This Agreement shall be governed by and construed in accordance with the laws of the '
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    r2 = s4.add_run('<Jurisdiction>')
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'
    r2.bold = True

    r3 = s4.add_run(
        ', without regard to its conflict of laws principles.'
    )
    r3.font.size = Pt(11)
    r3.font.name = 'Times New Roman'

    # Section 5
    doc.add_heading('5. Miscellaneous', level=2)
    s5 = doc.add_paragraph()
    r = s5.add_run(
        'This Agreement constitutes the entire agreement between the parties with respect to the '
        'subject matter hereof and supersedes all prior or contemporaneous agreements, understandings, '
        'negotiations, and discussions, whether oral or written. This Agreement may not be amended '
        'except by a written instrument signed by both parties.'
    )
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    # Signature block
    doc.add_paragraph()  # spacing
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(24)
    r = sig.add_run('AGREED AND ACCEPTED:')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    r.bold = True

    doc.add_paragraph()
    sig_line1 = doc.add_paragraph()
    r = sig_line1.add_run('_________________________')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    name_line = doc.add_paragraph()
    r = name_line.add_run('<PartyName>')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    r.bold = True

    comp_line = doc.add_paragraph()
    r = comp_line.add_run('<CompanyName>')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    doc.add_paragraph()
    sig_line2 = doc.add_paragraph()
    r = sig_line2.add_run('_________________________')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    disc_sig = doc.add_paragraph()
    r = disc_sig.add_run('Authorized Representative')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    disc_comp = doc.add_paragraph()
    r = disc_comp.add_run('Quantum Cipher Technologies, Inc.')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    doc.save(TEMPLATE_FILE)
    print(f'NDA template created: {TEMPLATE_FILE}')


def create_data_source():
    """Create CSV data source with 8 NDA party records."""
    with open(DATA_FILE, 'w', newline='') as f:
        writer = csv.DictWriter(f, fieldnames=['PartyName', 'CompanyName', 'EffectiveDate', 'Jurisdiction'])
        writer.writeheader()
        writer.writerows(NDA_PARTIES)
    print(f'Data source created: {DATA_FILE}')


def create_output_dir():
    """Create the output directory."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f'Output directory created: {OUTPUT_DIR}')


def main():
    create_nda_template()
    create_data_source()
    create_output_dir()

    # GUI-ready startup: open the NDA template in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{TEMPLATE_FILE}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with NDA_Template.docx')


main()
