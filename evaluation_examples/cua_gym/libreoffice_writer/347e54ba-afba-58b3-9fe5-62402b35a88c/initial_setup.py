"""
Initial Setup: Advanced textbook document with heading sections in Default Paragraph Style
Task ID: writer_struct_051
Domain: libreoffice_writer

Creates a ~20-page textbook with all section/chapter/part headings as
Default Paragraph Style (12pt). The agent must apply the proper heading hierarchy.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_051'
OUTPUT = f'{WORKDIR}/Desktop/advanced_textbook.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_body_paragraph(doc, text, num_sentences=4):
    """Add a body paragraph with the given text (Default Paragraph Style, 12pt)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    return para


def add_section_heading(doc, text):
    """Add a section heading as Default Paragraph Style 12pt (NOT a heading style)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    # Explicitly set to normal paragraph style — NOT heading
    para.style = doc.styles['Normal']
    return para


def create_initial():
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document()

    # Set default font size to 12pt
    style = doc.styles['Normal']
    style.font.size = Pt(12)

    # --- Title Page ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run("Advanced Studies in Systems Theory and Practice")
    run.bold = True
    run.font.size = Pt(18)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = subtitle_para.add_run("A Comprehensive Textbook for Graduate Students")
    run2.font.size = Pt(14)

    edition_para = doc.add_paragraph()
    edition_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = edition_para.add_run("Third Edition")
    run3.font.size = Pt(12)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = author_para.add_run("Dr. Eleanor Whitfield and Prof. James Hartley")
    run4.font.size = Pt(12)

    publisher_para = doc.add_paragraph()
    publisher_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run5 = publisher_para.add_run("Meridian Academic Press, 2024")
    run5.font.size = Pt(12)

    doc.add_page_break()

    # --- Preface ---
    preface_para = doc.add_paragraph()
    run = preface_para.add_run("Preface")
    run.bold = True
    run.font.size = Pt(14)

    add_body_paragraph(doc, (
        "This textbook has been designed for graduate-level courses in systems theory and applied methodology. "
        "Over the course of twenty chapters and four major parts, we guide the reader from foundational axioms "
        "through historical development, empirical case studies, and practical implementation strategies. "
        "The material has been refined over three editions based on extensive classroom feedback."
    ))

    add_body_paragraph(doc, (
        "Readers are expected to have a working knowledge of basic logic, introductory statistics, and "
        "undergraduate-level exposure to philosophy of science. Each part opens with a conceptual overview "
        "and closes with synthesis exercises designed to consolidate understanding across chapters. "
        "Instructors may find the supplemental materials available on our companion website helpful."
    ))

    add_body_paragraph(doc, (
        "We are deeply grateful to the many colleagues, reviewers, and students who contributed to this edition. "
        "Special thanks go to the research teams at the Institute for Complex Systems and to the editorial staff "
        "at Meridian Academic Press for their diligence and support throughout the revision process."
    ))

    doc.add_page_break()

    # ============================================================
    # PART I: Foundations  (should become Heading 1 — currently Normal 12pt)
    # ============================================================
    add_section_heading(doc, "Part I: Foundations")

    add_body_paragraph(doc, (
        "The first part of this textbook lays the conceptual groundwork for all subsequent analysis. "
        "We begin with an examination of core theoretical frameworks that have shaped the discipline, "
        "followed by a thorough grounding in the historical development of systems thinking. "
        "Readers are encouraged to engage critically with each chapter before moving forward."
    ))

    # Chapter 1: Theoretical Framework  (should become Heading 2 — currently Normal 12pt)
    add_section_heading(doc, "Chapter 1: Theoretical Framework")

    add_body_paragraph(doc, (
        "Systems theory emerged as a formalized discipline in the mid-twentieth century, drawing from "
        "biology, engineering, and cybernetics. This chapter surveys the principal intellectual currents "
        "that shaped its emergence: general systems theory, cybernetics, and information theory. "
        "We trace how these streams converged into a unified analytical methodology."
    ))

    add_body_paragraph(doc, (
        "Central to the theoretical framework is the concept of emergence — the idea that system-level "
        "properties cannot be fully predicted from the properties of individual components. "
        "This principle has profound implications for modeling, prediction, and intervention in complex systems. "
        "We examine both formal and informal treatments of emergence across scientific traditions."
    ))

    # Section 1.1: Core Axioms  (should become Heading 3 — currently Normal 12pt)
    add_section_heading(doc, "Section 1.1: Core Axioms")

    add_body_paragraph(doc, (
        "The axiomatic foundation of systems theory rests on a small set of core propositions. "
        "First, every system is composed of interrelated elements that form a coherent whole. "
        "Second, the relationships between elements are as important as the elements themselves. "
        "Third, every system is embedded within a broader environment with which it exchanges matter, energy, or information."
    ))

    add_body_paragraph(doc, (
        "These axioms generate a rich set of derived properties including feedback loops, equilibrium states, "
        "and phase transitions. The axiomatic approach allows us to apply systems concepts across wildly "
        "different empirical domains — from ecological networks to financial markets — while maintaining "
        "formal rigor and testability. Subsequent chapters will operationalize these axioms in detail."
    ))

    add_body_paragraph(doc, (
        "Philosophers of science have debated whether these axioms constitute genuine universal laws or "
        "merely useful heuristics. We take a pragmatist stance: the axioms are valuable insofar as they "
        "generate testable predictions and guide empirical inquiry. Their metaphysical status is secondary "
        "to their methodological utility."
    ))

    # Section 1.2: Derived Principles  (should become Heading 3 — currently Normal 12pt)
    add_section_heading(doc, "Section 1.2: Derived Principles")

    add_body_paragraph(doc, (
        "From the core axioms, several important derived principles follow with logical necessity. "
        "The principle of equifinality holds that a system can reach the same final state from different "
        "initial conditions and by different paths. This challenges classical deterministic models and "
        "has important implications for causal inference in complex systems research."
    ))

    add_body_paragraph(doc, (
        "The principle of requisite variety, originally proposed by Ashby, states that an effective "
        "controller must have at least as many response states as there are disturbance states in the "
        "environment it seeks to regulate. This principle has been applied in management science, "
        "ecological resilience theory, and robust control engineering."
    ))

    add_body_paragraph(doc, (
        "Homeostasis and negative feedback constitute another set of derived principles with wide "
        "empirical application. Living organisms, organizations, and engineered systems all employ "
        "feedback mechanisms to maintain stable operating states. The formal analysis of such mechanisms "
        "is a core competency for systems practitioners."
    ))

    doc.add_page_break()

    # Chapter 2: Historical Context  (should become Heading 2 — currently Normal 12pt)
    add_section_heading(doc, "Chapter 2: Historical Context")

    add_body_paragraph(doc, (
        "Understanding the historical development of systems thinking illuminates not only where the field "
        "stands today but also why certain methodological commitments have become canonical. "
        "This chapter traces the intellectual genealogy of systems theory from its origins in the early "
        "twentieth century through the current era of computational modeling."
    ))

    add_body_paragraph(doc, (
        "The Vienna Circle's logical positivism and the Gestalt psychologists' holism represent two "
        "pre-war intellectual currents that, in different ways, prefigured modern systems thinking. "
        "The wartime development of operations research and the post-war blossoming of cybernetics "
        "under Norbert Wiener and his collaborators marked a decisive turning point."
    ))

    add_body_paragraph(doc, (
        "By the 1960s and 1970s, systems thinking had penetrated management science through the work of "
        "Jay Forrester and his colleagues at MIT, who developed system dynamics as a methodology for "
        "simulating the behavior of complex sociotechnical systems. Forrester's World Dynamics models "
        "sparked intense controversy and ultimately contributed to the global sustainability discourse."
    ))

    add_body_paragraph(doc, (
        "The Santa Fe Institute, founded in 1984, became a crucible for complexity science — a close "
        "intellectual relative of systems theory that brought physicists, economists, and biologists "
        "together to study adaptive complex systems. Their work on nonlinear dynamics, agent-based "
        "modeling, and network theory has significantly influenced contemporary systems research."
    ))

    add_body_paragraph(doc, (
        "Today the field is undergoing a further transformation driven by big data analytics, machine "
        "learning, and advances in network science. These tools offer unprecedented capacity for "
        "empirical systems analysis at scale, though they also raise new challenges for theory "
        "construction and causal explanation."
    ))

    doc.add_page_break()

    # ============================================================
    # PART II: Applications  (should become Heading 1 — currently Normal 12pt)
    # ============================================================
    add_section_heading(doc, "Part II: Applications")

    add_body_paragraph(doc, (
        "The second part of the textbook turns from theoretical foundations to empirical applications. "
        "We examine how systems concepts have been applied in legal reasoning, business strategy, "
        "and organizational design. The goal is to demonstrate the practical power of a systems "
        "perspective for analyzing and solving real-world problems."
    ))

    # Chapter 3: Case Studies  (should become Heading 2 — currently Normal 12pt)
    add_section_heading(doc, "Chapter 3: Case Studies")

    add_body_paragraph(doc, (
        "Case studies occupy a privileged position in applied systems research. They allow the analyst "
        "to examine system dynamics in their full complexity, without the simplifying assumptions "
        "required by formal models. This chapter presents two clusters of cases — one from legal "
        "practice, one from business — to illustrate key systems concepts in action."
    ))

    add_body_paragraph(doc, (
        "Each case study is structured around a central paradox or tension: a situation in which "
        "linear, parts-focused thinking fails to capture system-level dynamics. By working through "
        "the cases systematically, students develop intuition for recognizing systems patterns "
        "in complex empirical situations."
    ))

    # Section 3.1: Legal Cases  (should become Heading 3 — currently Normal 12pt)
    add_section_heading(doc, "Section 3.1: Legal Cases")

    add_body_paragraph(doc, (
        "The legal system is itself a complex adaptive system characterized by feedback loops between "
        "precedent, legislative response, and social norms. Our first case examines the unintended "
        "consequences of the Three Strikes mandatory sentencing legislation enacted in several U.S. "
        "states during the 1990s. A simple deterrence model predicted declining recidivism rates; "
        "the actual outcomes were considerably more complex."
    ))

    add_body_paragraph(doc, (
        "Prison populations expanded dramatically, disproportionately affecting communities of color "
        "and straining state budgets to the point where other public services were cut. Paradoxically, "
        "some research suggests that the law may have increased certain categories of violent crime, "
        "as offenders facing life sentences had incentives to eliminate witnesses. These feedback "
        "effects were invisible to the linear deterrence model."
    ))

    add_body_paragraph(doc, (
        "Our second legal case examines the development of environmental standing doctrine in U.S. "
        "administrative law. The expansion of citizen suit provisions in environmental statutes "
        "created positive feedback loops between advocacy organizations, agency rulemaking, and "
        "judicial interpretation. Mapping these loops helps explain the nonlinear trajectory of "
        "environmental law over the past fifty years."
    ))

    # Section 3.2: Business Cases  (should become Heading 3 — currently Normal 12pt)
    add_section_heading(doc, "Section 3.2: Business Cases")

    add_body_paragraph(doc, (
        "Business organizations exhibit classic systems behaviors: goal-seeking feedback, oscillation, "
        "growth and collapse archetypes, and emergent strategies. Our first business case examines "
        "the supply chain dynamics of the semiconductor industry, whose pronounced boom-bust cycles "
        "illustrate the bullwhip effect — the amplification of demand variability as one moves "
        "upstream in a supply chain."
    ))

    add_body_paragraph(doc, (
        "The semiconductor case is particularly instructive because the industry's capital-intensive "
        "nature means that supply adjustments are slow and costly, while demand fluctuations can be "
        "rapid and severe. The result is a structural tendency toward oscillation that has persisted "
        "despite decades of attempts at supply chain rationalization. Systems modeling of the industry's "
        "dynamics reveals the structural origins of this pattern."
    ))

    add_body_paragraph(doc, (
        "Our second business case examines Kodak's failure to adapt to digital photography. "
        "Kodak actually invented the digital camera in 1975, yet systematically underinvested in "
        "digital technology for decades. A systems analysis reveals the reinforcing loops that locked "
        "the company into its film-based business model and the balancing loops that prevented "
        "adequate response to the digital disruption signal until it was too late."
    ))

    doc.add_page_break()

    # Chapter 4: Implementation  (should become Heading 2 — currently Normal 12pt)
    add_section_heading(doc, "Chapter 4: Implementation")

    add_body_paragraph(doc, (
        "Theory and case analysis are necessary but not sufficient for effective systems practice. "
        "This chapter addresses the practical challenge of implementing systems thinking within "
        "organizations and policy processes. We draw on lessons from decades of applied systems "
        "work to identify success factors and common failure modes."
    ))

    add_body_paragraph(doc, (
        "Successful implementation typically requires three elements: a committed leadership sponsor, "
        "a multidisciplinary team capable of integrating quantitative and qualitative insights, and "
        "a deliberate process for moving from systems diagnosis to intervention design. "
        "Organizations that rush to solutions without adequate diagnostic work consistently "
        "encounter the problem of 'fixes that fail' — interventions that address symptoms rather "
        "than underlying structural causes."
    ))

    add_body_paragraph(doc, (
        "The systems dynamics methodology developed at MIT provides one well-tested implementation "
        "framework. Key steps include problem articulation, dynamic hypothesis development, model "
        "formulation and testing, policy design and evaluation, and organizational learning. "
        "Each step is iterative: findings at later stages routinely require revisiting earlier stages. "
        "The process is more spiral than linear."
    ))

    add_body_paragraph(doc, (
        "Soft systems methodology (SSM), developed by Peter Checkland, offers an alternative "
        "framework better suited to situations where stakeholders disagree about the nature of the "
        "problem. SSM uses structured facilitation processes and rich pictures to elicit multiple "
        "perspectives, compare them systematically, and identify feasible and desirable changes "
        "that accommodate legitimate differences in value and priority."
    ))

    add_body_paragraph(doc, (
        "Whichever framework is employed, effective implementation requires sustained attention to "
        "organizational culture and power dynamics. Systems insights that threaten established "
        "interests or challenge dominant mental models will encounter resistance regardless of "
        "their technical merits. Building systems literacy across an organization — not just among "
        "specialists — is a long-term investment that pays dividends in adaptive capacity."
    ))

    doc.add_page_break()

    # --- Bibliography ---
    bib_para = doc.add_paragraph()
    run = bib_para.add_run("Bibliography")
    run.bold = True
    run.font.size = Pt(14)

    references = [
        "Ashby, W. R. (1956). An Introduction to Cybernetics. London: Chapman and Hall.",
        "Beer, S. (1972). Brain of the Firm. London: Allen Lane.",
        "Checkland, P. (1981). Systems Thinking, Systems Practice. Chichester: Wiley.",
        "Forrester, J. W. (1961). Industrial Dynamics. Cambridge, MA: MIT Press.",
        "Meadows, D. H. (2008). Thinking in Systems: A Primer. White River Junction, VT: Chelsea Green.",
        "Senge, P. (1990). The Fifth Discipline: The Art and Practice of the Learning Organization. New York: Doubleday.",
        "Simon, H. A. (1962). The Architecture of Complexity. Proceedings of the American Philosophical Society, 106(6), 467-482.",
        "von Bertalanffy, L. (1968). General System Theory. New York: George Braziller.",
        "Wiener, N. (1948). Cybernetics: Or Control and Communication in the Animal and the Machine. Cambridge, MA: MIT Press.",
        "Sterman, J. D. (2000). Business Dynamics: Systems Thinking and Modeling for a Complex World. Boston: McGraw-Hill.",
    ]

    for ref in references:
        ref_para = doc.add_paragraph()
        run = ref_para.add_run(ref)
        run.font.size = Pt(11)
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_page_break()

    # --- Index ---
    index_para = doc.add_paragraph()
    run = index_para.add_run("Index")
    run.bold = True
    run.font.size = Pt(14)

    index_entries = [
        "Adaptive systems, 12, 45, 89, 134",
        "Ashby, William Ross, 23, 67",
        "Bullwhip effect, 178, 183",
        "Checkland, Peter, 201, 215",
        "Complexity science, 45, 56, 89",
        "Cybernetics, 15, 22, 67, 89",
        "Emergence, 8, 34, 56, 78",
        "Equifinality, 44, 67",
        "Feedback loops, 20, 34, 45, 145, 178",
        "Forrester, Jay, 55, 89, 201",
        "Homeostasis, 45, 67, 89",
        "Kodak Corporation, 185, 189",
        "Mental models, 215, 220",
        "Negative feedback, 45, 67, 134",
        "Positive feedback, 145, 155, 165",
        "Requisite variety, 44, 67",
        "Santa Fe Institute, 57, 78",
        "Semiconductor industry, 178",
        "Senge, Peter, 201, 220",
        "System dynamics, 55, 89, 201",
        "Three Strikes legislation, 155, 163",
        "von Bertalanffy, Ludwig, 15, 22, 34",
        "Wiener, Norbert, 22, 34, 55",
    ]

    for entry in index_entries:
        entry_para = doc.add_paragraph()
        run = entry_para.add_run(entry)
        run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
