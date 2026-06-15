"""
Reward Script: Build a comprehensive product catalog document from a large Calc dataset
Task ID: osworld_multi_apps_doc_calc_to_writer_011
Domain: libreoffice_writer
Scoring:
  - Component 1: product_catalog.odt exists in Documents (precondition gate)
  - Component 2: Cover page text 'Product Catalog 2024' present (0.15)
  - Component 3: 5 category H1 headings present (0.25)
  - Component 4: 5 product tables with correct column headers and row counts (0.25)
  - Component 5: Pricing Summary table with correct columns and 5 rows (0.20)
  - Component 6: Footer contains page number field (0.15)
  Total: 1.0
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_doc_calc_to_writer_011'
FILE_PATH = os.path.join(WORKDIR, 'Documents', 'product_catalog.odt')

# Expected categories in order
EXPECTED_CATEGORIES = ['Electronics', 'Clothing', 'Food', 'Books']
# Home Goods may appear as 'Home Goods' or 'Home_Goods'
HOME_GOODS_VARIANTS = ['Home Goods', 'Home_Goods', 'HomeGoods']

# Expected row counts per category (data rows, excluding header)
EXPECTED_ROW_COUNTS = [10, 8, 7, 8, 7]

# Expected product table headers
EXPECTED_PRODUCT_HEADERS = ['Product_ID', 'Name', 'Description', 'Price', 'Stock',
                             'Product ID', 'Unit_Price', 'Stock_Count']

# Expected summary columns
EXPECTED_SUMMARY_COLS = ['Category', 'Avg', 'Min', 'Max', 'Total']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: File must exist
    if not os.path.exists(file_path):
        print('FAIL: product_catalog.odt not found at %s' % file_path)
        print('REWARD: 0.0')
        return 0.0

    # Load the document
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print('CRITICAL: Cannot load file %s: %s' % (file_path, e))
        print('REWARD: 0.0')
        return 0.0

    print('File loaded successfully')
    print('Paragraphs: %d, Tables: %d, Sections: %d' % (
        len(doc.paragraphs), len(doc.tables), len(doc.sections)))

    # Component 1: Cover page contains 'Product Catalog 2024' (0.15 points)
    try:
        cover_found = False
        for para in doc.paragraphs:
            if 'Product Catalog 2024' in para.text or 'product catalog 2024' in para.text.lower():
                cover_found = True
                break
        if cover_found:
            print('PASS: Component 1 — Cover page text "Product Catalog 2024" found (0.15 pts)')
            total_score += 0.15
        else:
            all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            print('FAIL: Component 1 — "Product Catalog 2024" not found. First paras: %s' % all_texts[:5])
    except Exception as e:
        print('ERROR: Component 1 — %s' % e)

    # Component 2: 5 category H1 headings (0.25 points)
    # Each heading is worth 0.05 points
    try:
        h1_texts = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if 'Heading 1' in style_name or style_name == 'Heading 1':
                h1_texts.append(para.text.strip())

        print('Found H1 headings: %s' % h1_texts)

        # Check for each expected category
        found_categories = 0
        for cat in EXPECTED_CATEGORIES:
            found = any(cat.lower() in h.lower() for h in h1_texts)
            if found:
                found_categories += 1
                print('  PASS: H1 heading for "%s" found' % cat)
            else:
                print('  FAIL: H1 heading for "%s" not found' % cat)

        # Check for Home Goods (may appear in various forms)
        home_goods_found = any(
            any(variant.lower() in h.lower() for variant in HOME_GOODS_VARIANTS)
            for h in h1_texts
        )
        if home_goods_found:
            found_categories += 1
            print('  PASS: H1 heading for "Home Goods" found')
        else:
            print('  FAIL: H1 heading for "Home Goods/Home_Goods" not found')

        heading_score = round(found_categories * 0.05, 2)
        if found_categories == 5:
            print('PASS: Component 2 — All 5 category H1 headings found (0.25 pts)')
        else:
            print('PASS partial: Component 2 — %d/5 category H1 headings found (%.2f pts)' % (
                found_categories, heading_score))
        total_score += heading_score
    except Exception as e:
        print('ERROR: Component 2 — %s' % e)

    # Component 3: 5 product tables with correct structure (0.25 points)
    # Each table is worth 0.05 points: 0.025 for headers, 0.025 for row count
    try:
        product_tables = []
        summary_table = None

        # The summary table is the one with 'Category' as first header
        for table in doc.tables:
            if table.rows:
                first_row = [c.text.strip() for c in table.rows[0].cells]
                if first_row and 'Category' in first_row[0]:
                    summary_table = table
                else:
                    product_tables.append(table)

        print('Found %d product tables and %s summary table' % (
            len(product_tables), 'a' if summary_table else 'no'))

        table_score = 0.0
        for t_idx, table in enumerate(product_tables[:5]):
            first_row = [c.text.strip() for c in table.rows[0].cells]
            data_rows = len(table.rows) - 1

            # Check headers contain expected columns
            has_correct_headers = (
                any('id' in h.lower() for h in first_row) and
                any('name' in h.lower() for h in first_row) and
                any('desc' in h.lower() for h in first_row) and
                any('price' in h.lower() for h in first_row) and
                any('stock' in h.lower() for h in first_row)
            )

            # Check expected row count
            expected_count = EXPECTED_ROW_COUNTS[t_idx] if t_idx < len(EXPECTED_ROW_COUNTS) else 0
            has_correct_rows = (data_rows == expected_count)

            sub_score = 0.0
            if has_correct_headers:
                sub_score += 0.025
                print('  PASS: Table %d has correct column headers' % t_idx)
            else:
                print('  FAIL: Table %d headers wrong: %s' % (t_idx, first_row))

            if has_correct_rows:
                sub_score += 0.025
                print('  PASS: Table %d has correct row count (%d)' % (t_idx, data_rows))
            else:
                print('  FAIL: Table %d row count %d, expected %d' % (t_idx, data_rows, expected_count))

            table_score += sub_score

        if len(product_tables) < 5:
            print('FAIL: Component 3 — Only %d product tables found, expected 5' % len(product_tables))
        else:
            print('Component 3 — product table score: %.3f / 0.25' % table_score)
        total_score += table_score
    except Exception as e:
        print('ERROR: Component 3 — %s' % e)

    # Component 4: Pricing Summary table (0.20 points)
    try:
        if summary_table is not None:
            # Check summary table headers
            header_row = [c.text.strip() for c in summary_table.rows[0].cells]
            print('Summary table headers: %s' % header_row)

            has_category_col = any('category' in h.lower() for h in header_row)
            has_avg_col = any('avg' in h.lower() for h in header_row)
            has_min_col = any('min' in h.lower() for h in header_row)
            has_max_col = any('max' in h.lower() for h in header_row)
            has_total_col = any('total' in h.lower() or 'item' in h.lower() for h in header_row)

            # Check data rows (should have 5 categories)
            data_rows = len(summary_table.rows) - 1

            summary_score = 0.0
            if has_category_col and has_avg_col and has_min_col and has_max_col and has_total_col:
                summary_score += 0.10
                print('PASS: Component 4 — Summary table has correct columns (0.10 pts)')
            else:
                print('FAIL: Component 4 — Summary table missing columns. Has: %s' % header_row)

            if data_rows == 5:
                summary_score += 0.10
                print('PASS: Component 4 — Summary table has 5 data rows (0.10 pts)')
            else:
                print('FAIL: Component 4 — Summary table has %d data rows, expected 5' % data_rows)

            total_score += summary_score
        else:
            print('FAIL: Component 4 — No pricing summary table found')
    except Exception as e:
        print('ERROR: Component 4 — %s' % e)

    # Component 5: Footer contains page numbers (0.15 points)
    try:
        section = doc.sections[0]
        footer = section.footer
        footer_has_page = False

        for para in footer.paragraphs:
            xml_str = para._element.xml
            # Check for PAGE field code (standard field code for page numbers)
            if 'PAGE' in xml_str or 'fldChar' in xml_str or 'instrText' in xml_str:
                footer_has_page = True
                break
            # Also check if footer has any page-related text
            if 'page' in para.text.lower() or any(c.isdigit() for c in para.text):
                footer_has_page = True
                break

        if footer_has_page:
            print('PASS: Component 5 — Footer contains page number field (0.15 pts)')
            total_score += 0.15
        else:
            print('FAIL: Component 5 — Footer does not contain page number field')
            # Check footer content for debugging
            for para in footer.paragraphs:
                print('  Footer para text: %r' % para.text)
    except Exception as e:
        print('ERROR: Component 5 — %s' % e)

    final_score = min(round(total_score, 4), 1.0)
    print('\nScore: %.4f/1.0' % total_score)
    print('REWARD: %s' % final_score)
    return final_score


# Entry point
if not os.path.exists(FILE_PATH):
    print('File not found: %s' % FILE_PATH)
    print('REWARD: 0.0')
else:
    verify_task(FILE_PATH)
