"""
Reward Script: Format academic paper in IEEE two-column format
Task ID: writer_pd_032
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Page margins — 1.78cm top/bottom, 1.65cm left/right
  Component 2 (0.20): Title formatting — 24pt bold centered
  Component 3 (0.15): Author line — 11pt centered
  Component 4 (0.20): Abstract — 'Abstract' bold 9pt centered, body 9pt italic
  Component 5 (0.20): Two-column layout for body text (section with cols=2)
"""

import os

from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_032'
WNS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def emu_to_cm(emu):
    """Convert EMU to centimeters."""
    if emu is None:
        return 0
    return emu / 914400 * 2.54


def get_col_count(section):
    """Get column count from section XML."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    cols_elem = section._sectPr.find('.//w:cols', ns)
    if cols_elem is not None:
        num = cols_elem.get(WNS + 'num')
        if num is not None:
            return int(num)
    return 1


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Page margins — 1.78cm top/bottom, 1.65cm left/right (0.25 points)
    # Check ALL sections have the correct margins (tolerance 0.05cm)
    try:
        all_margins_ok = True
        margin_details = []
        for i, s in enumerate(doc.sections):
            tm = emu_to_cm(s.top_margin)
            bm = emu_to_cm(s.bottom_margin)
            lm = emu_to_cm(s.left_margin)
            rm = emu_to_cm(s.right_margin)
            margin_details.append(f"Sec{i}: top={tm:.3f} bot={bm:.3f} left={lm:.3f} right={rm:.3f}")

            if abs(tm - 1.78) > 0.05 or abs(bm - 1.78) > 0.05:
                all_margins_ok = False
            if abs(lm - 1.65) > 0.05 or abs(rm - 1.65) > 0.05:
                all_margins_ok = False

        if all_margins_ok and len(doc.sections) >= 1:
            print(f"PASS: Component 1 — Margins correct: {margin_details} (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Margins incorrect: {margin_details}. Expected ~1.78cm top/bot, ~1.65cm left/right")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Title formatting — first paragraph is 24pt bold centered (0.20 points)
    try:
        title_para = doc.paragraphs[0]
        title_centered = str(title_para.paragraph_format.alignment) in ['CENTER (1)', '1']
        # Check runs for 24pt bold
        title_bold = False
        title_24pt = False
        for r in title_para.runs:
            if r.font.bold is True:
                title_bold = True
            if r.font.size and abs(r.font.size.pt - 24.0) < 0.5:
                title_24pt = True

        if title_centered and title_bold and title_24pt:
            print(f"PASS: Component 2 — Title is 24pt bold centered (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — Title: centered={title_centered}, bold={title_bold}, 24pt={title_24pt}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Author line — 11pt centered (0.15 points)
    try:
        author_para = doc.paragraphs[1]
        author_centered = str(author_para.paragraph_format.alignment) in ['CENTER (1)', '1']
        author_11pt = False
        for r in author_para.runs:
            if r.font.size and abs(r.font.size.pt - 11.0) < 0.5:
                author_11pt = True

        if author_centered and author_11pt:
            print(f"PASS: Component 3 — Author line is 11pt centered (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 3 — Author: centered={author_centered}, 11pt={author_11pt}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Abstract section — 'Abstract' heading bold 9pt, body 9pt italic (0.20 points)
    try:
        # Find the paragraph with text 'Abstract'
        abstract_heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().lower() == 'abstract':
                abstract_heading_idx = i
                break

        if abstract_heading_idx is None:
            print("FAIL: Component 4 — 'Abstract' heading not found")
        else:
            abs_para = doc.paragraphs[abstract_heading_idx]
            # Check Abstract heading: bold and 9pt
            abs_bold = False
            abs_9pt = False
            for r in abs_para.runs:
                if r.font.bold is True:
                    abs_bold = True
                if r.font.size and abs(r.font.size.pt - 9.0) < 0.5:
                    abs_9pt = True

            # Check abstract body (next non-empty paragraph): 9pt italic
            abs_body_italic = False
            abs_body_9pt = False
            for j in range(abstract_heading_idx + 1, min(abstract_heading_idx + 3, len(doc.paragraphs))):
                bp = doc.paragraphs[j]
                if bp.text.strip():
                    for r in bp.runs:
                        if r.font.italic is True:
                            abs_body_italic = True
                        if r.font.size and abs(r.font.size.pt - 9.0) < 0.5:
                            abs_body_9pt = True
                    break

            sub_score = 0.0
            if abs_bold and abs_9pt:
                sub_score += 0.10
                print(f"  PASS: Abstract heading is bold 9pt")
            else:
                print(f"  FAIL: Abstract heading bold={abs_bold}, 9pt={abs_9pt}")

            if abs_body_italic and abs_body_9pt:
                sub_score += 0.10
                print(f"  PASS: Abstract body is 9pt italic")
            else:
                print(f"  FAIL: Abstract body italic={abs_body_italic}, 9pt={abs_body_9pt}")

            if sub_score > 0:
                print(f"PASS: Component 4 — Abstract formatting ({sub_score} pts)")
                total_score += sub_score
            else:
                print(f"FAIL: Component 4 — Abstract formatting incorrect")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Two-column layout for body text (0.20 points)
    # Golden has >=2 sections: one single-column (for title/abstract) and one 2-column (for body)
    try:
        has_two_col_section = False
        for s in doc.sections:
            if get_col_count(s) >= 2:
                has_two_col_section = True
                break

        if has_two_col_section and len(doc.sections) >= 2:
            print(f"PASS: Component 5 — Two-column layout found with {len(doc.sections)} sections (0.20 pts)")
            total_score += 0.20
        elif has_two_col_section:
            # Has two columns but only one section (partial credit)
            print(f"PARTIAL: Component 5 — Two-column layout found but only {len(doc.sections)} section(s) (0.10 pts)")
            total_score += 0.10
        else:
            col_counts = [get_col_count(s) for s in doc.sections]
            print(f"FAIL: Component 5 — No two-column section found. Column counts: {col_counts}")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
