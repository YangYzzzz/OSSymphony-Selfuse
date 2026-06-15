"""
Initial Setup: Format academic paper in IEEE two-column format
Task ID: writer_pd_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_pd_032'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Default margins (2.54 cm) - explicitly set to ensure defaults
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ---- Title (unformatted - default style, NOT bold, NOT 24pt, NOT centered) ----
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(
        "Deep Reinforcement Learning for Autonomous Navigation in Complex Urban Environments"
    )
    title_run.font.size = Pt(11)  # default size
    title_run.font.name = "Calibri"

    # ---- Author names (unformatted - default style) ----
    author_para = doc.add_paragraph()
    author_run = author_para.add_run(
        "Wei Zhang, Sarah Chen, Marcus Rivera, and Priya Patel"
    )
    author_run.font.size = Pt(11)
    author_run.font.name = "Calibri"

    # ---- Affiliation ----
    affil_para = doc.add_paragraph()
    affil_run = affil_para.add_run(
        "Department of Computer Science, Stanford University, Stanford, CA 94305"
    )
    affil_run.font.size = Pt(11)
    affil_run.font.name = "Calibri"

    # ---- Abstract (unformatted - no bold heading, no 9pt italic) ----
    abs_heading_para = doc.add_paragraph()
    abs_heading_run = abs_heading_para.add_run("Abstract")
    abs_heading_run.font.size = Pt(11)
    abs_heading_run.font.name = "Calibri"

    abstract_text = (
        "This paper presents a novel deep reinforcement learning framework for autonomous "
        "vehicle navigation in dense urban settings. Our approach combines a hierarchical "
        "policy architecture with a multi-modal perception system that fuses LiDAR point clouds, "
        "camera imagery, and HD map data. We introduce a curriculum-based training strategy "
        "that progressively increases environmental complexity, enabling the agent to develop "
        "robust navigation behaviors. Extensive experiments on the CARLA simulator and real-world "
        "test tracks demonstrate that our method achieves a 94.7% success rate in complex "
        "intersection scenarios, outperforming existing approaches by 12.3 percentage points. "
        "We further show that the learned policies transfer effectively to previously unseen "
        "urban environments with minimal fine-tuning."
    )
    abs_body_para = doc.add_paragraph()
    abs_body_run = abs_body_para.add_run(abstract_text)
    abs_body_run.font.size = Pt(11)
    abs_body_run.font.name = "Calibri"

    # ---- Body text sections (all single-column, default formatting) ----

    # Section 1: Introduction
    doc.add_paragraph()
    s1_heading = doc.add_paragraph()
    s1_run = s1_heading.add_run("I. Introduction")
    s1_run.font.size = Pt(11)
    s1_run.font.name = "Calibri"

    intro_paragraphs = [
        "Autonomous navigation in urban environments remains one of the most challenging "
        "problems in robotics and artificial intelligence. The complexity arises from the "
        "need to handle dynamic obstacles, unpredictable pedestrian behavior, varying traffic "
        "conditions, and diverse road geometries. Traditional rule-based approaches, while "
        "effective in structured highway scenarios, struggle to generalize across the wide "
        "variety of situations encountered in city driving.",

        "Recent advances in deep reinforcement learning (DRL) have shown promising results "
        "in sequential decision-making tasks, including game playing, robot manipulation, and "
        "simplified driving scenarios. However, applying DRL to real-world urban navigation "
        "presents unique challenges: the state space is extremely high-dimensional, the action "
        "space must be continuous and precise, and the reward signal is sparse and delayed.",

        "In this work, we propose HierNav, a hierarchical deep reinforcement learning framework "
        "specifically designed for urban autonomous navigation. Our architecture separates the "
        "navigation task into strategic route planning and tactical maneuver execution, enabling "
        "more efficient learning and better generalization. The strategic planner operates on a "
        "graph representation of the road network, while the tactical controller processes raw "
        "sensor data to execute safe, comfortable driving maneuvers.",

        "Our key contributions are as follows: (1) We introduce a hierarchical policy architecture "
        "that decouples strategic planning from tactical execution, reducing the effective "
        "dimensionality of the learning problem. (2) We propose a multi-modal perception module "
        "that fuses LiDAR, camera, and map data through cross-attention mechanisms. (3) We "
        "develop a curriculum learning strategy that enables stable training in progressively "
        "complex environments. (4) We demonstrate state-of-the-art performance on both simulated "
        "and real-world benchmarks."
    ]
    for text in intro_paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Section 2: Related Work
    doc.add_paragraph()
    s2_heading = doc.add_paragraph()
    s2_run = s2_heading.add_run("II. Related Work")
    s2_run.font.size = Pt(11)
    s2_run.font.name = "Calibri"

    related_paragraphs = [
        "A. Deep Reinforcement Learning for Driving",

        "The application of DRL to autonomous driving has been explored extensively in recent "
        "years. Dosovitskiy et al. demonstrated end-to-end driving using conditional imitation "
        "learning in the CARLA simulator. Liang et al. combined model-based and model-free RL "
        "approaches for urban driving, achieving improved sample efficiency. Chen et al. proposed "
        "a latent space RL method that disentangles scene understanding from decision making. "
        "Our work differs from these approaches by introducing a hierarchical decomposition that "
        "explicitly separates strategic from tactical reasoning.",

        "B. Multi-Modal Sensor Fusion",

        "Effective perception in urban environments requires integrating information from multiple "
        "sensor modalities. PointPillars and VoxelNet established efficient representations for "
        "LiDAR processing, while BEVFormer and LSS pioneered bird's-eye-view representations "
        "from camera inputs. TransFuser demonstrated the benefits of combining camera and LiDAR "
        "features through transformer-based attention mechanisms. Our perception module builds "
        "upon these foundations but introduces a novel cross-modal attention scheme optimized "
        "for the downstream RL objective.",

        "C. Hierarchical Reinforcement Learning",

        "Hierarchical RL methods decompose complex tasks into manageable sub-problems. The "
        "options framework provides a theoretical foundation for temporal abstraction in RL. "
        "Feudal Networks and HIRO demonstrated practical implementations of hierarchical "
        "policies for continuous control tasks. In the driving domain, hierarchy has been "
        "applied to separate lane-level decisions from continuous control. Our approach extends "
        "this paradigm by integrating map-level strategic planning with sensor-driven tactical "
        "control, bridging the gap between route planning and reactive driving."
    ]
    for text in related_paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Section 3: Method
    doc.add_paragraph()
    s3_heading = doc.add_paragraph()
    s3_run = s3_heading.add_run("III. Methodology")
    s3_run.font.size = Pt(11)
    s3_run.font.name = "Calibri"

    method_paragraphs = [
        "A. System Overview",

        "The HierNav framework consists of three main components: (1) a multi-modal perception "
        "module that processes raw sensor data into a unified feature representation, (2) a "
        "strategic planner that selects waypoints along the route based on traffic conditions "
        "and map information, and (3) a tactical controller that generates continuous vehicle "
        "control commands to reach the selected waypoints while avoiding obstacles.",

        "B. Multi-Modal Perception Module",

        "Our perception module takes as input a point cloud from a 64-beam LiDAR sensor, "
        "RGB images from six surround-view cameras, and a local HD map patch centered on the "
        "ego vehicle. The LiDAR point cloud is encoded using a PointPillars-style backbone "
        "that produces a 200x200 bird's-eye-view (BEV) feature grid at 0.25m resolution. "
        "Camera images are processed through a ResNet-50 backbone followed by a lift-splat-shoot "
        "module to produce camera BEV features at the same resolution.",

        "The fusion of LiDAR and camera BEV features is performed through a cross-attention "
        "transformer with 4 attention heads and 3 layers. The HD map features, encoded as a "
        "rasterized multi-channel image, are concatenated channel-wise after a lightweight "
        "convolutional encoder. The final fused representation is a 256-dimensional feature "
        "vector for each spatial location in the BEV grid.",

        "C. Strategic Planner",

        "The strategic planner operates on a graph representation of the road network, where "
        "nodes represent intersections and lane merge/split points, and edges represent road "
        "segments. At each decision point, the planner selects the next intermediate waypoint "
        "from a set of candidates derived from the global route. The planner's policy is "
        "parameterized as a graph attention network (GAT) that takes as input node features "
        "(traffic density, signal state, road type) and edge features (segment length, speed "
        "limit, lane count).",

        "The planner is trained using Proximal Policy Optimization (PPO) with a reward function "
        "that balances route progress, estimated travel time, and safety margins. The planner "
        "updates at a frequency of 2 Hz, providing waypoints approximately 50-100 meters ahead "
        "of the ego vehicle.",

        "D. Tactical Controller",

        "The tactical controller receives the selected waypoint from the strategic planner "
        "along with the fused perception features and generates continuous control commands: "
        "steering angle, throttle, and brake. The controller is parameterized as a recurrent "
        "neural network (GRU) with 3 layers and 512 hidden units, operating at 10 Hz.",

        "The tactical controller is trained using Soft Actor-Critic (SAC) with a composite "
        "reward function incorporating: waypoint tracking error (distance and heading), "
        "comfort metrics (jerk and lateral acceleration), collision penalties, and traffic "
        "rule compliance. We use a hindsight experience replay buffer to improve sample "
        "efficiency in sparse-reward scenarios.",

        "E. Curriculum Learning Strategy",

        "Training directly in complex urban environments leads to poor convergence due to the "
        "sparse reward signal and high-dimensional state space. We propose a four-stage "
        "curriculum that progressively introduces complexity: Stage 1 covers straight road "
        "following with no traffic. Stage 2 adds curved roads and static obstacles. Stage 3 "
        "introduces dynamic traffic participants. Stage 4 presents full urban scenarios with "
        "intersections, traffic signals, and pedestrians. Advancement between stages is based "
        "on achieving a success rate threshold of 90% at the current stage."
    ]
    for text in method_paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Section 4: Experiments
    doc.add_paragraph()
    s4_heading = doc.add_paragraph()
    s4_run = s4_heading.add_run("IV. Experiments")
    s4_run.font.size = Pt(11)
    s4_run.font.name = "Calibri"

    exp_paragraphs = [
        "A. Experimental Setup",

        "We evaluate HierNav on two benchmarks: the CARLA NoCrash benchmark and a custom "
        "real-world test track in Sunnyvale, California. The CARLA benchmark includes 25 "
        "predefined routes across 4 weather conditions and 2 traffic density levels. Our "
        "real-world test track covers 15 km of urban roads with 8 signalized intersections, "
        "3 roundabouts, and mixed pedestrian-vehicle traffic zones.",

        "All models are trained on 8 NVIDIA A100 GPUs for 5 million environment steps in "
        "CARLA, requiring approximately 72 hours of wall-clock time. The tactical controller "
        "processes observations at 10 Hz, while the strategic planner updates at 2 Hz. We "
        "use a replay buffer capacity of 1 million transitions and a batch size of 256.",

        "B. Baseline Comparisons",

        "We compare HierNav against the following baselines: (1) TransFuser, an end-to-end "
        "imitation learning approach using transformer-based sensor fusion; (2) MILE, a "
        "model-based RL method with learned world models; (3) TCP, a trajectory-conditioned "
        "policy that separates planning from control; and (4) InterFuser, which uses a "
        "safety-aware interpretable sensor fusion approach.",

        "C. Results on CARLA NoCrash Benchmark",

        "Table I presents the results on the CARLA NoCrash benchmark. HierNav achieves a "
        "94.7% success rate in the Dense Traffic condition, outperforming the next best method "
        "(InterFuser) by 12.3 percentage points. In the Regular Traffic condition, HierNav "
        "achieves 97.2% success rate. The improvement is particularly pronounced in complex "
        "intersection scenarios, where the hierarchical decomposition allows the strategic "
        "planner to anticipate traffic patterns and the tactical controller to execute precise "
        "maneuvers.",

        "The infraction rate analysis reveals that HierNav significantly reduces collision "
        "frequency compared to baselines: 0.23 collisions per km versus 0.41 for InterFuser "
        "and 0.67 for TransFuser. Red light violations are also reduced by 45% compared to "
        "the best baseline, attributable to the strategic planner's explicit modeling of "
        "traffic signal states.",

        "D. Real-World Evaluation",

        "On our Sunnyvale test track, HierNav completes 18 out of 20 evaluation routes "
        "without any safety-critical interventions. The two incomplete routes involved "
        "construction zones with temporary lane closures not present in the training "
        "distribution. Average travel time is within 8% of human driver performance, "
        "and passenger comfort metrics (measured by RMS lateral acceleration) are "
        "0.14 m/s squared compared to 0.18 m/s squared for the human baseline."
    ]
    for text in exp_paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # Section 5: Conclusion
    doc.add_paragraph()
    s5_heading = doc.add_paragraph()
    s5_run = s5_heading.add_run("V. Conclusion")
    s5_run.font.size = Pt(11)
    s5_run.font.name = "Calibri"

    conclusion_paragraphs = [
        "We have presented HierNav, a hierarchical deep reinforcement learning framework "
        "for autonomous navigation in complex urban environments. By decomposing the navigation "
        "task into strategic planning and tactical control, our approach achieves state-of-the-art "
        "performance on both simulated and real-world benchmarks. The curriculum learning strategy "
        "enables stable training, while the multi-modal perception module provides robust "
        "environmental understanding.",

        "Future work will focus on extending HierNav to handle more diverse weather conditions, "
        "incorporating vehicle-to-vehicle communication data, and developing safe exploration "
        "strategies for continued learning during deployment. We believe hierarchical approaches "
        "represent a promising direction for scaling autonomous navigation to the full complexity "
        "of real-world urban driving.",
    ]
    for text in conclusion_paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    # References
    doc.add_paragraph()
    ref_heading = doc.add_paragraph()
    ref_run = ref_heading.add_run("References")
    ref_run.font.size = Pt(11)
    ref_run.font.name = "Calibri"

    references = [
        "[1] A. Dosovitskiy, G. Ros, F. Codevilla, A. Lopez, and V. Koltun, \"CARLA: An open urban driving simulator,\" in CoRL, 2017.",
        "[2] X. Liang, T. Wang, L. Yang, and E. Xing, \"CIRL: Controllable imitative reinforcement learning for vision-based self-driving,\" in ECCV, 2018.",
        "[3] D. Chen, B. Zhou, V. Koltun, and P. Krahenbuhl, \"Learning by cheating,\" in CoRL, 2020.",
        "[4] A. Lang, S. Vora, H. Caesar, L. Zhou, J. Yang, and O. Beijbom, \"PointPillars: Fast encoders for object detection from point clouds,\" in CVPR, 2019.",
        "[5] Z. Li, W. Wang, H. Li, E. Xie, C. Sima, T. Lu, Y. Qiao, and J. Dai, \"BEVFormer: Learning bird's-eye-view representation from multi-camera images via spatiotemporal transformers,\" in ECCV, 2022.",
        "[6] P. Chitta, A. Prakash, B. Jaeger, Z. Yu, K. Renz, and A. Geiger, \"TransFuser: Imitation with transformer-based sensor fusion for autonomous driving,\" IEEE T-PAMI, 2023.",
        "[7] H. Shao, L. Wang, R. Chen, H. Li, and Y. Liu, \"Safety-enhanced autonomous driving using interpretable sensor fusion transformer,\" in CoRL, 2023.",
        "[8] R. Sutton, D. Precup, and S. Singh, \"Between MDPs and semi-MDPs: A framework for temporal abstraction in reinforcement learning,\" Artificial Intelligence, 1999.",
        "[9] A. Vezhnevets, S. Osindero, T. Schaul, N. Heess, M. Jaderberg, D. Silver, and K. Kavukcuoglu, \"FeUdal networks for hierarchical reinforcement learning,\" in ICML, 2017.",
        "[10] O. Nachum, S. Gu, H. Lee, and S. Levine, \"Data-efficient hierarchical reinforcement learning,\" in NeurIPS, 2018.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
