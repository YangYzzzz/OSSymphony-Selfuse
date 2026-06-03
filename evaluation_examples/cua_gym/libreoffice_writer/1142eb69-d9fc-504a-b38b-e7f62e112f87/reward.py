"""
Reward Script: Apply widow and orphan control to the entire document
Task ID: writer_rd_070
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Normal/Default paragraph style has widowControl enabled
  Component 2 (0.4): >= 80% of paragraphs have widowControl ON (explicit or inherited)
  Component 3 (0.2): No paragraphs have widowControl explicitly disabled (val=0)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_070'


def persist_app_state(domain: str):
    """Save any unsaved changes in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify widow and orphan control is applied to the entire document.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not available")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    wns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Component 1: Normal/Default paragraph style has widowControl enabled (0.4 points)
    # This is the primary change: the Default Paragraph Style must have widow control ON.
    try:
        styles_element = doc.styles.element
        normal_style_wc = None
        for style in styles_element.findall('.//w:style', ns):
            style_id = style.attrib.get(f'{wns}styleId', '')
            style_type = style.attrib.get(f'{wns}type', '')
            is_default = style.attrib.get(f'{wns}default', '0')
            # Find the Normal/default paragraph style
            if style_type == 'paragraph' and (style_id == 'Normal' or is_default == '1'):
                pPr = style.find('w:pPr', ns)
                if pPr is not None:
                    wc = pPr.find('w:widowControl', ns)
                    if wc is not None:
                        val = wc.attrib.get(f'{wns}val', '1')
                        # In OOXML, val="1" or val="true" or absent val all mean ON
                        # val="0" or val="false" means OFF
                        if val in ('1', 'true'):
                            normal_style_wc = True
                        else:
                            normal_style_wc = False
                    else:
                        # widowControl absent in style pPr means default (ON in OOXML spec)
                        normal_style_wc = True
                else:
                    # No pPr means all defaults, widowControl defaults to ON
                    normal_style_wc = True
                break

        if normal_style_wc is True:
            print(f"PASS: Component 1 — Normal style has widowControl enabled (0.4 pts)")
            total_score += 0.4
        elif normal_style_wc is False:
            print(f"FAIL: Component 1 — Normal style has widowControl explicitly DISABLED")
        else:
            print(f"FAIL: Component 1 — Could not find Normal/default paragraph style")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: >= 80% of paragraphs have widowControl ON (0.4 points)
    # The task says "entire document" — we check that the vast majority of paragraphs
    # have widowControl enabled, either explicitly or by inheriting from the style.
    try:
        total_paras = len(doc.paragraphs)
        wc_on_count = 0
        wc_off_count = 0
        wc_inherit_count = 0  # no explicit setting, inherits from style

        for para in doc.paragraphs:
            pPr = para._element.find('w:pPr', ns)
            if pPr is not None:
                wc = pPr.find('w:widowControl', ns)
                if wc is not None:
                    val = wc.attrib.get(f'{wns}val', '1')
                    if val in ('1', 'true'):
                        wc_on_count += 1
                    else:
                        wc_off_count += 1
                else:
                    wc_inherit_count += 1
            else:
                wc_inherit_count += 1

        # Paragraphs that inherit: if Normal style is ON, these are effectively ON
        effective_on = wc_on_count
        if normal_style_wc is True:
            effective_on += wc_inherit_count

        if total_paras > 0:
            ratio = effective_on / total_paras
        else:
            ratio = 0.0

        print(f"  Paragraph stats: total={total_paras}, explicit_on={wc_on_count}, "
              f"explicit_off={wc_off_count}, inherit={wc_inherit_count}, "
              f"effective_on={effective_on}, ratio={ratio:.2%}")

        if ratio >= 0.80:
            print(f"PASS: Component 2 — {ratio:.0%} of paragraphs have widowControl ON (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Only {ratio:.0%} of paragraphs have widowControl ON (need >= 80%)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: No paragraphs have widowControl explicitly disabled (0.2 points)
    # For the task to be fully complete, no paragraph should override with val="0".
    try:
        if wc_off_count == 0:
            print(f"PASS: Component 3 — No paragraphs have widowControl explicitly disabled (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — {wc_off_count} paragraphs have widowControl explicitly disabled")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
