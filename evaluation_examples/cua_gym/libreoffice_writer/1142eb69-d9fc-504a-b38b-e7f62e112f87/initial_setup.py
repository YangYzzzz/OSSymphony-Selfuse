"""
Initial Setup: Novel manuscript with widow/orphan control disabled
Task ID: writer_rd_070
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_070'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def disable_widow_orphan(para):
    """Explicitly disable widow and orphan control on a paragraph."""
    pPr = para._element.get_or_add_pPr()
    widow = pPr.makeelement(qn('w:widowControl'), {qn('w:val'): '0'})
    pPr.append(widow)


def create_initial():
    doc = Document()

    # Set up page margins for standard manuscript format
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Disable widow/orphan control on the default paragraph style
    default_style = doc.styles['Normal']
    pf_xml = default_style.element.get_or_add_pPr()
    wc = pf_xml.makeelement(qn('w:widowControl'), {qn('w:val'): '0'})
    pf_xml.append(wc)

    # Novel title page
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(200)
    run = title_para.add_run("THE MERIDIAN CHRONICLES")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("A Novel")
    run.italic = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(40)
    run = author.add_run("by Elena Vasquez")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_page_break()

    # Novel content - enough to fill ~25 pages with realistic prose
    chapters = {
        "Chapter 1: The Arrival": [
            "The train pulled into Meridian Station at exactly half past six, just as the autumn sun began to paint the sky in shades of amber and crimson. Clara Whitfield stepped onto the platform with nothing but a leather satchel and the crumpled letter from her late grandmother that had brought her across three state lines to this peculiar town she had never visited before.",
            "The station itself was a relic of another era. Iron columns supported a glass canopy that had long since lost several panes, allowing rain and leaves to accumulate in drifts along the wooden benches. A single attendant sat behind a brass-grilled window, reading a newspaper that appeared to be several days old. He did not look up when Clara approached.",
            "\"Excuse me,\" she said, her voice echoing slightly in the cavernous space. \"I'm looking for Hawthorn Lane. Could you point me in the right direction?\"",
            "The attendant lowered his newspaper slowly, revealing a weathered face with deep-set eyes that regarded her with mild suspicion. \"Hawthorn Lane,\" he repeated, as though testing the words for some hidden meaning. \"That's up past the old mill, beyond the covered bridge. About a twenty-minute walk if you keep a good pace. But I wouldn't recommend going after dark.\"",
            "Clara thanked him and stepped outside into the cool evening air. The town of Meridian spread before her like a watercolor painting left out in the rain, its colors bleeding softly into one another. Victorian houses lined the main street, their gingerbread trim and wraparound porches speaking of prosperity that had come and gone. Most of the shops were closed, their windows dark, though a warm glow emanated from what appeared to be a tavern at the far end of the street.",
            "She consulted the letter again, though she had memorized its contents long ago. Her grandmother, Margaret Whitfield, had written it three months before her death, in a hand that trembled but remained legible. The property at 14 Hawthorn Lane was Clara's now, along with whatever secrets it contained. Margaret had been characteristically cryptic about the details.",
            "\"You will understand when you see it,\" the letter read. \"Some things cannot be explained in writing. They must be experienced. Go to the house, Clara. Go before winter comes. The garden will show you what I could never tell you in life.\"",
            "Clara had dismissed it initially as the ramblings of an elderly woman whose mind had begun to wander. But the lawyer's letter that followed, with its official seals and witnessed signatures, confirmed that the property was real, the inheritance legitimate, and that certain conditions were attached to its transfer. She must take up residence within six months of Margaret's passing or the house would revert to the Meridian Historical Trust.",
            "The walk through town was pleasant enough, despite the growing darkness. Clara passed a general store with barrels of apples stacked outside, a tiny post office with a flag that had seen better days, and a white-clapboard church whose steeple disappeared into the gathering clouds. A few people nodded to her as she passed, their expressions curious but not unfriendly.",
            "The covered bridge appeared just as the attendant had described, spanning a creek that ran silver in the fading light. Clara's footsteps echoed on the wooden planks as she crossed, and she paused halfway to look down at the water. Something caught her eye, a flash of color beneath the surface, but when she looked again there was nothing but the smooth flow of the current over moss-covered stones.",
            "Beyond the bridge, the road narrowed to a dirt path lined with ancient oaks whose branches interlocked overhead like the ribs of a cathedral. Hawthorn Lane branched off to the left, marked by a rusted sign half-hidden in a tangle of wild roses. Number fourteen was the last house on the lane, set back from the road behind an overgrown garden that even in its neglected state hinted at former grandeur.",
            "The house itself was larger than Clara had expected, a three-story Victorian with a turret room, bay windows, and a porch that wrapped around three sides. Paint peeled from the clapboards in long strips, and several shutters hung at odd angles, but the structure appeared fundamentally sound. The roof was intact, the chimney straight, and the front steps, though weathered, held her weight without complaint.",
        ],
        "Chapter 2: The Garden": [
            "Clara woke to the sound of birdsong, a chorus so elaborate and sustained that for a moment she thought someone had left a radio playing. She lay still in the narrow bed she had found in the second-floor bedroom, staring up at the ceiling where a water stain had created a pattern that looked remarkably like a map of an unknown coastline.",
            "The house had been less forbidding by the light of the battery-powered lantern she had brought. The rooms were dusty but furnished, as though someone had simply walked out one day and never returned. A kitchen with a wood-burning stove, a parlor with bookshelves stretching to the ceiling, a dining room with a table set for six. Everything covered in a fine layer of dust and silence.",
            "She dressed quickly and went downstairs to explore by daylight. The kitchen window faced east, and morning sun streamed through the glass, illuminating motes of dust that danced in the warm air. Clara found a jar of instant coffee in the cupboard that was only slightly past its expiration date and heated water on the stove, which mercifully still worked.",
            "Coffee in hand, she opened the back door and stepped into the garden. What she saw made her stop mid-stride, the mug halfway to her lips, steam curling upward and mingling with her suspended breath.",
            "The garden was enormous, stretching at least two acres behind the house and bordered on three sides by a stone wall covered in climbing ivy. But it was not the size that astonished her. It was the plants themselves. Even in their wild, untended state, they were extraordinary. Colors she had never seen in any nursery or botanical garden blazed from every corner. Flowers with petals that seemed to shift hue as you watched, leaves with patterns that looked almost deliberately geometric, vines that spiraled in perfect mathematical ratios.",
            "In the center of the garden stood a fountain, dry now, its basin cracked but still beautiful. Around it, arranged in concentric circles, were raised beds that had once been carefully tended. Clara could see the bones of the original design beneath the overgrowth, a mandala-like pattern that radiated outward from the fountain in ever-widening rings.",
            "She set down her coffee and walked along what had once been a gravel path, now mostly reclaimed by moss and creeping thyme. Each section of the garden seemed to have its own microclimate. Near the south wall, where the stone absorbed and radiated the sun's warmth, Mediterranean herbs thrived, their fragrance almost overwhelming. In a shaded corner to the north, she found ferns and mosses of astonishing variety, some with fronds so delicate they were nearly translucent.",
            "But it was the section nearest the fountain that drew her attention most. Here, the plants were unlike anything she had encountered in her decade of horticultural study. A shrub with leaves that appeared to be a deep indigo blue, not the blue-green of certain sages, but a true, saturated blue that looked almost artificial. Beside it, a climbing rose whose flowers graduated from deep crimson at the base to pure white at the tips, each petal a perfect gradient.",
            "Clara knelt beside a patch of ground cover and touched one of the tiny leaves. It was warm, noticeably warmer than the surrounding soil, as though the plant generated its own heat. She pulled her hand back in surprise, then reached out again, more slowly. Yes, the warmth was unmistakable.",
            "\"Remarkable, isn't it?\" said a voice from behind her.",
            "Clara spun around, nearly losing her balance. An elderly man stood at the garden gate, leaning on a walking stick carved from a single piece of dark wood. He wore a tweed jacket despite the mild morning and had the kind of face that seemed perpetually amused by something only he could see.",
            "\"I'm sorry, I didn't mean to startle you,\" he said, raising one hand in apology. \"I'm Thomas Blackwood. I was Margaret's neighbor, and her friend. I've been keeping an eye on the place since she passed. You must be Clara.\"",
            "Clara stood and brushed the dirt from her knees. \"You knew my grandmother?\"",
            "\"For forty years. She was the most extraordinary woman I've ever known. And this garden was her masterpiece.\" He gestured with his walking stick at the wild profusion of plants. \"Though I suspect it doesn't look quite the way she left it.\"",
        ],
        "Chapter 3: The Journals": [
            "Thomas Blackwood stayed for lunch, which Clara improvised from supplies she had brought, crackers, cheese, dried fruit, and the instant coffee that was rapidly becoming a staple. They sat on the porch, plates balanced on the wide railing, and he told her about Margaret.",
            "\"She came to Meridian in nineteen seventy-two,\" he said, breaking a cracker into precise halves. \"Bought this house for almost nothing. It had been empty for years, decades really, and the locals considered it a lost cause. But Margaret saw something in it that others didn't. Or perhaps she simply saw what it could become.\"",
            "\"Was she always interested in gardening?\" Clara asked.",
            "Thomas smiled. \"That's rather like asking whether Mozart was interested in music. Your grandmother didn't garden, Clara. She conducted symphonies with soil and seeds. But I'm getting ahead of myself. You should read her journals first. They're in the turret room, unless someone has moved them, which I doubt. No one has been inside since Margaret's last visit here, and that was nearly a year before she died.\"",
            "After Thomas left, promising to return the next day with supplies from town, Clara climbed the narrow stairs to the turret room. It was circular, as she had expected, with windows on all sides that offered a panoramic view of the garden and the surrounding landscape. In the distance, she could see the covered bridge and, beyond it, the steeple of the white church.",
            "The room was furnished as a study. A desk occupied the center, positioned to take advantage of the light from every direction. Bookshelves curved along the walls, filled with volumes on botany, chemistry, folklore, and what appeared to be several languages Clara didn't recognize. And on the desk, arranged in chronological order, was a series of leather-bound journals spanning from nineteen seventy-three to the previous year.",
            "Clara picked up the first journal and opened it carefully. Margaret's handwriting was younger here, more confident, the letters formed with the precision of someone who had been taught penmanship as a serious discipline. The first entry was dated March fifteenth, nineteen seventy-three.",
            "\"The soil here is unlike anything I've tested before. pH varies by location, sometimes within a matter of feet. The northeast corner reads six-point-two, while the south wall area shows nearly eight. This should be impossible in a space this small without deliberate amendment, and yet the existing plant life suggests these conditions have been stable for a very long time.\"",
            "Clara read on, increasingly absorbed. Margaret's journals were not mere gardening notes. They were scientific observations of the highest order, meticulous, systematic, and increasingly astonished. Over the months and years that followed her arrival, Margaret had catalogued dozens of plant specimens that did not appear in any published taxonomy. She had measured growth rates that exceeded known biological limits. She had documented properties that bordered on the inexplicable.",
            "\"The blue-leafed shrub (designation MG-017) has no known relatives in any family I can identify,\" read an entry from nineteen seventy-eight. \"DNA analysis at the university lab returned results that the technician described as 'corrupted,' but I suspect the sample was perfectly valid. This plant simply does not fit into our existing classification system. It may require a new one.\"",
            "By the time the light began to fail, Clara had read through the first three journals and was beginning the fourth. Her coffee had gone cold hours ago. Her cheese and crackers sat untouched. She was aware of a growing excitement, the kind she had not felt since her graduate school days when every experiment seemed to open a door onto a new room full of questions.",
            "Margaret had not been a gardener at all, Clara realized. She had been a researcher, and this property was her laboratory. The garden was not a collection of pretty flowers. It was a carefully controlled experiment that had been running for fifty years. And the results, if Margaret's observations were to be believed, challenged everything Clara thought she knew about plant biology.",
            "She closed the journal and looked out the turret window at the garden below, now silver and shadow in the moonlight. The blue-leafed shrub was visible even in the darkness, its leaves seeming to glow faintly with their own light. Clara watched it for a long moment, then turned back to the journals.",
        ],
        "Chapter 4: The Experiment": [
            "The next morning brought rain, a soft persistent drizzle that turned the garden paths to mud and drummed against the turret windows with hypnotic regularity. Clara had fallen asleep at the desk sometime after midnight and woke with a stiff neck and a journal page imprinted on her cheek.",
            "She made coffee and returned to reading. The fourth journal covered nineteen eighty through eighty-two, a period Margaret described as \"the breakthrough years.\" The entries became longer, more detailed, and increasingly excited, though Margaret's scientific precision never wavered.",
            "\"I have confirmed the preliminary results,\" Margaret wrote in June of nineteen eighty. \"MG-017 and MG-023 are not merely unusual specimens. They are evidence of a fundamentally different approach to photosynthesis. Where conventional plants convert light to chemical energy through chlorophyll, these specimens appear to utilize a pigment I have designated meridianin, which absorbs across a much broader spectrum. The implications are staggering.\"",
            "Clara set down the journal and stared out at the rain. She was a botanist herself, with a PhD from Cornell and twelve years of post-doctoral research. She knew what Margaret was describing was impossible by every standard of modern plant science. And yet the garden outside was full of plants that should not exist.",
            "She thought about calling her colleague David Chen at the university. David was a plant biochemist, cautious by nature, but open to new ideas. He would approach Margaret's claims with skepticism but also with the intellectual honesty to test them properly. But something held her back. The letter had said to experience this, not to share it. At least not yet.",
            "The rain continued through the afternoon. Clara explored the house more thoroughly, finding a cellar stocked with preserved seeds in labeled glass jars, a potting shed attached to the south side of the house with tools and soil amendments in neatly organized rows, and a small laboratory in what had once been a pantry, complete with a microscope, slides, and reagent bottles.",
            "In the laboratory, she found a locked cabinet. The key was in Margaret's desk drawer, hanging from a chain with a tag that read simply \"Careful.\" Inside the cabinet were three glass containers, each holding a small amount of crystalline substance. The labels read \"Meridianin A,\" \"Meridianin B,\" and \"Meridianin C.\" Beside them was a notebook, smaller than the journals, bound in red leather.",
            "The red notebook was Margaret's private research log, distinct from the journals. It contained detailed protocols for extracting, purifying, and testing the meridianin compounds. It also contained warnings, underlined and annotated in red ink, about dosage, handling, and something Margaret called \"cascade effects.\"",
            "\"Under no circumstances should meridianin be applied to conventional plant stock without careful titration,\" one warning read. \"The growth response is exponential, not linear, and the window between therapeutic and destructive doses is alarmingly narrow. I lost an entire bed of heritage roses in nineteen seventy-nine by miscalculating the dilution by a factor of ten. The plants grew to six feet in three days, then collapsed into undifferentiated cellular mass. It was not a pleasant sight.\"",
            "Clara closed the red notebook and sat quietly in the dim pantry laboratory, listening to the rain on the roof. She was beginning to understand why Margaret had kept this work private. The scientific establishment would have dismissed it as fantasy. The commercial world would have exploited it without understanding or caution. And the government, well, Clara could imagine several agencies that would have been very interested in a compound that could make things grow exponentially.",
            "Her grandmother had chosen a different path. She had spent fifty years studying the phenomenon in secret, documenting everything with rigorous scientific methodology, and waiting for someone she trusted to continue the work. Someone with the training to understand it and the wisdom to handle it responsibly.",
            "Clara looked at the three glass containers in the locked cabinet and felt the weight of Margaret's trust settle onto her shoulders like a physical thing. She locked the cabinet, pocketed the key, and went upstairs to continue reading the journals.",
        ],
        "Chapter 5: The Neighbors": [
            "By the end of her first week in Meridian, Clara had established a routine. Mornings were for the garden, where she began the slow work of clearing overgrowth and mapping the layout of Margaret's plantings. Afternoons were for the journals, which she read methodically, taking her own notes in a fresh notebook she had purchased from the general store in town. Evenings were for cooking, reading fiction from Margaret's extensive library, and trying not to think too hard about what she was learning.",
            "Thomas Blackwood visited every other day, always arriving precisely at ten in the morning with some small offering, eggs from his chickens, bread from the bakery in town, once a jar of honey from his own bees. He was a retired professor of folklore and mythology, which struck Clara as either coincidental or deeply significant, given the nature of Margaret's research.",
            "\"Did my grandmother ever talk to you about the garden?\" Clara asked one morning, as they walked along the stone wall that bordered the eastern edge of the property.",
            "Thomas was quiet for a moment. \"She told me enough,\" he said finally. \"Enough to know that this place is special. Enough to know that it needs protecting. Not the kind of protecting that involves locks and fences, though those have their place. The kind that involves understanding.\"",
            "\"Understanding what, exactly?\"",
            "He stopped and turned to face her, his expression more serious than she had yet seen. \"Clara, this land has been different for as long as anyone can remember. The indigenous people who lived here before European settlement considered it sacred ground. They didn't farm it or build on it. They visited it for ceremonies, and their medicine people came here to gather plants that grew nowhere else.\"",
            "\"How do you know this?\"",
            "\"Because I spent thirty years researching it. Margaret and I collaborated, in our way. She studied the science. I studied the stories. Different lenses on the same phenomenon.\" He gestured at the garden with his walking stick. \"Every culture that has encountered this place has recognized it as extraordinary. The question is why.\"",
            "Clara was about to respond when a movement near the south wall caught her eye. A woman was standing just outside the garden gate, watching them. She was tall and thin, with iron-gray hair pulled back in a severe bun, and she wore a long dark coat despite the warm morning.",
            "\"That's Evelyn March,\" Thomas said quietly, following Clara's gaze. \"She lives in the house on the other side of the lane. She and Margaret did not get along.\"",
            "The woman held Clara's gaze for a long moment, then turned and walked away without speaking. Clara felt an unexpected chill that had nothing to do with the temperature.",
            "\"What was their disagreement about?\" she asked.",
            "Thomas sighed. \"Evelyn believes the property should belong to the Historical Trust. She's been on the board for twenty years and has never made a secret of her desire to acquire this land. She sees it as a public asset that Margaret selfishly kept private.\"",
            "\"And what do you think?\"",
            "\"I think Margaret knew exactly what she was doing, and she had very good reasons for keeping this place out of institutional hands. Reasons that you will come to understand as you read further in those journals.\"",
            "Clara looked at the spot where Evelyn March had stood. The woman had left no footprints in the soft earth, which struck her as odd. But then Thomas was asking about her progress with the garden restoration, and she let the observation slip away, filing it in the growing cabinet of things she would need to think about later.",
        ],
        "Chapter 6: The Discovery": [
            "On the fourteenth day, Clara found the underground chamber. She had been clearing dead vegetation from around the base of the fountain when her trowel struck something metallic. Careful excavation revealed a brass ring set into a stone slab, green with age but still solid. The slab lifted with surprising ease, revealing stone steps descending into darkness.",
            "Clara's first instinct was caution. She went back to the house for a flashlight, a rope, and her phone, which she set to record video. Then she descended the steps, counting as she went. Twelve steps down, each carved from a single block of granite, worn smooth by centuries of use.",
            "The chamber at the bottom was roughly circular, perhaps twenty feet in diameter, with a domed ceiling just high enough for Clara to stand upright. The walls were lined with stone shelves on which sat dozens of clay vessels, each sealed with wax. The air was cool and dry, with a faintly mineral scent that reminded Clara of caves she had visited in her college geology classes.",
            "In the center of the chamber, directly below where the fountain stood above, was a shallow basin carved into the floor. It was filled with soil, dark, rich-looking soil that seemed almost to pulse with vitality. Clara knelt beside it and pressed her fingers into the earth. It was warm, almost hot, and it tingled against her skin.",
            "She sat back on her heels and swept the flashlight beam around the chamber. On the wall opposite the stairs, she noticed something she had missed at first: a series of symbols carved into the stone, running in a band from floor to ceiling. They were not any script she recognized, neither Latin nor Greek nor any indigenous writing system she had studied. They were organic in form, resembling roots and branches and seed pods, as though someone had tried to write using the alphabet of plants.",
            "Clara photographed everything methodically, then returned to the surface. She sealed the opening and covered it with the vegetation she had cleared, not out of secrecy but out of a sudden protective instinct she couldn't fully articulate. Then she went to the turret room and began searching the journals for any mention of an underground chamber.",
            "She found it in journal seven, dated nineteen eighty-nine. \"Finally located the source,\" Margaret wrote, her handwriting betraying an excitement that had broken through her usual scientific restraint. \"The chamber beneath the fountain is the origin point. The soil there is not soil in any conventional sense. It is a living matrix, a substrate that has been accumulating and evolving for what may be thousands of years. Every extraordinary property of this garden traces back to that chamber.\"",
            "Clara read on, her pulse quickening. Margaret had spent the next decade studying the chamber soil, running every test she could devise with the equipment available to her. The results were consistent and extraordinary. The soil contained microorganisms that did not match any known species. It produced organic compounds that had no analogues in existing chemistry databases. And it had a measurable effect on any plant material placed in contact with it.",
            "\"I have tested sixty-seven species of conventional plants,\" Margaret wrote in nineteen ninety-three. \"In every case, extended contact with the chamber soil produces significant modifications. Enhanced growth rate, novel pigmentation, increased resistance to disease and environmental stress, and in thirty-eight cases, the development of entirely new biochemical pathways. The plants do not merely grow better. They become something new.\"",
            "Clara set down the journal and looked out at the garden. The rain had stopped, and late afternoon sun painted everything in gold. She could see the fountain from the turret window, its cracked basin innocent and ordinary-looking. But she knew now what lay beneath it, and the knowledge changed everything.",
            "She thought about David Chen again, and about the locked cabinet with its three glass containers. She thought about Evelyn March, watching from beyond the gate. She thought about Margaret, alone with this secret for fifty years, documenting everything, trusting no one until the very end when she had written that final letter to a granddaughter she barely knew.",
        ],
    }

    for chapter_title, paragraphs in chapters.items():
        # Chapter heading
        heading = doc.add_heading(chapter_title, level=1)
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)
        disable_widow_orphan(heading)

        # Chapter paragraphs
        for text in paragraphs:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            run = para.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            disable_widow_orphan(para)

        # Add spacing between chapters
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(24)
        disable_widow_orphan(spacer)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
