"""
FINAL REWARD SCRIPT - SUCCESS
Task: I’m in the middle of typing some meeting notes and now I need a small table right at the current cursor position—exactly 2 columns and 5 rows, with the first row set as a header so it stays separate when I sort or format. How do I pop that in quickly in LibreOffice Writer?
Generated: 2025-09-10 12:04:40
Status: success
Model: azure-o3
Total Steps: 5
"""

from docx import Document
from docx.oxml.ns import qn
import os


def verify_libreoffice_writer_table_task(file_path):
    """Reward script for LibreOffice-Writer task.

    Verifies that the document at *file_path* contains:
        1. EXACTLY one (or more) table with 2 columns and 5 rows (0.4 pts)
        2. First row of that table flagged as a header row (0.3 pts)
        3. Table appears after at least one non-empty paragraph (simulates insertion at current
           cursor in the middle of notes) (0.2 pts)
        4. Original meeting-notes text still present (keywords “Meeting Notes” & “Agenda item”) (0.1 pt)

    Progressive scoring is applied; only a fully‐correct solution yields 1.0.
    """

    print(f"Starting verification for file: {file_path}\n")
    total_score = 0.0
    MAX_SCORE = 1.0

    # Weights for each requirement
    WEIGHTS = {
        "table_size": 0.4,
        "header_row": 0.3,
        "placement": 0.2,
        "content_preserved": 0.1,
    }

    # ---------- 1. Load the document (prerequisite – no points) ----------
    if not os.path.exists(file_path):
        print("✗ File not found – task failed.")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
        print(f"✓ Document loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} table(s) found.\n")
    except Exception as e:
        print(f"✗ Could not open document: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ---------- 2. Locate a 5×2 table ----------
    target_table = None
    for idx, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if tbl.rows else 0
        print(f"  Table {idx + 1}: {rows} rows x {cols} columns")
        if rows == 5 and cols == 2 and target_table is None:
            target_table = tbl
    if target_table is not None:
        total_score += WEIGHTS["table_size"]
        print(f"✓ Correct-sized 5×2 table found (+{WEIGHTS['table_size']})\n")
    else:
        print("✗ No 5×2 table detected\n")

    # ---------- 3. Verify header row flag ----------
    if target_table is not None:
        first_row = target_table.rows[0]
        header_flag = False
        tr = first_row._tr  # CT_Row element
        if tr.trPr is not None:
            for child in tr.trPr:
                if child.tag == qn("w:tblHeader"):
                    header_flag = True
                    break
        if header_flag:
            total_score += WEIGHTS["header_row"]
            print(f"✓ First row marked as header (+{WEIGHTS['header_row']})\n")
        else:
            print("✗ First row NOT flagged as a header\n")

    # ---------- 4. Verify placement (table after existing text) ----------
    body_elems = list(doc.element.body)
    paragraphs_before_table = 0
    saw_table = False
    for el in body_elems:
        local_name = el.tag.split('}')[-1]
        if local_name == 'tbl':
            saw_table = True
            break
        if local_name == 'p':
            # Count only non-empty paragraphs
            texts = [node.text for node in el.iter() if node.tag.split('}')[-1] == 't' and node.text]
            if ''.join(texts).strip():
                paragraphs_before_table += 1
    if saw_table and paragraphs_before_table > 0:
        total_score += WEIGHTS["placement"]
        print(f"✓ Table placement verified after {paragraphs_before_table} paragraph(s) (+{WEIGHTS['placement']})\n")
    else:
        print("✗ Unable to verify correct table placement\n")

    # ---------- 5. Verify original meeting notes still present ----------
    expected_phrases = ["meeting notes", "agenda item"]
    all_text = ' '.join(p.text.lower() for p in doc.paragraphs if p.text)
    if all(phrase in all_text for phrase in expected_phrases):
        total_score += WEIGHTS["content_preserved"]
        print(f"✓ Meeting notes content preserved (+{WEIGHTS['content_preserved']})\n")
    else:
        print("✗ Required meeting-notes text missing or altered\n")

    # ---------- Final score ----------
    final_score = 1.0 if abs(total_score - MAX_SCORE) < 1e-6 else round(total_score, 4)
    print(f"Total score: {final_score}/{MAX_SCORE}")
    print(f"REWARD: {final_score}")
    return final_score


# ------------------------------------------------------------------
# Execute verification (path is fixed by platform for the candidate)
# ------------------------------------------------------------------
if __name__ == "__main__":
    TEST_FILE = "/home/user/im_in_the_middle_of_typing_some_meeting_notes_and_now_i_need_a_small_table_right_at_the_current_curs.docx"
    verify_libreoffice_writer_table_task(TEST_FILE)
