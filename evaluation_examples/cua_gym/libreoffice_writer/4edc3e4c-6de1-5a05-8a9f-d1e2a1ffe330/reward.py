"""
Reward Script: Insert custom TOC entries for 'Glossary of Terms' and 'List of Abbreviations'
Task ID: writer_struct_038
Domain: libreoffice_writer
Scoring:
  Component 1: TC field code for 'Glossary of Terms' at level 1 in body paragraph (0.4 pts)
  Component 2: TC field code for 'List of Abbreviations' at level 1 in body paragraph (0.4 pts)
  Component 3: TOC section updated — both entries appear in TOC text (0.2 pts)
  Total: 1.0
"""

import os
import re

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_038'


def find_tc_field_in_paragraph(para, expected_title):
    """
    Check if a paragraph contains a TC (Table of Contents entry) field code
    that references expected_title at level 1.

    A TC field looks like:
      <w:instrText> TC "Glossary of Terms" \l 1 </w:instrText>
    in the paragraph's XML.
    """
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    instr_elements = para._element.findall('.//w:instrText', ns)
    for instr in instr_elements:
        if instr.text:
            text = instr.text.strip()
            # Must start with TC, contain the expected title, and specify level 1
            if text.startswith('TC') and expected_title in text:
                # Check level: \l 1 or \l "1"
                if re.search(r'\\l\s+"?1"?', text):
                    return True, text
    return False, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires:
    1. The paragraph 'Glossary of Terms' has a TC field mark at level 1
    2. The paragraph 'List of Abbreviations' has a TC field mark at level 1
    3. The TOC section (near start of doc) contains both entries
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file " + file_path + ": " + str(e))
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: TC field code in 'Glossary of Terms' body paragraph (0.4 pts)
    # This MUST fail on initial_env (no TC fields present) and pass on golden_env
    # -----------------------------------------------------------------------
    try:
        glossary_tc_found = False
        glossary_tc_text = None
        for para in doc.paragraphs:
            para_text = para.text.strip()
            # Find the body paragraph (not the TOC line which has a tab+page number)
            if 'Glossary of Terms' in para_text and '\t' not in para_text:
                found, tc_text = find_tc_field_in_paragraph(para, 'Glossary of Terms')
                if found:
                    glossary_tc_found = True
                    glossary_tc_text = tc_text
                    break

        if glossary_tc_found:
            print("PASS: Component 1 — TC field found for 'Glossary of Terms' at level 1 (0.4 pts)")
            print("  instrText: " + str(glossary_tc_text))
            total_score += 0.4
        else:
            print("FAIL: Component 1 — No TC field at level 1 found in 'Glossary of Terms' paragraph")
    except Exception as e:
        print("ERROR: Component 1 — " + str(e))

    # -----------------------------------------------------------------------
    # Component 2: TC field code in 'List of Abbreviations' body paragraph (0.4 pts)
    # This MUST fail on initial_env (no TC fields present) and pass on golden_env
    # -----------------------------------------------------------------------
    try:
        abbrev_tc_found = False
        abbrev_tc_text = None
        for para in doc.paragraphs:
            para_text = para.text.strip()
            # Find the body paragraph (not the TOC line which has a tab+page number)
            if 'List of Abbreviations' in para_text and '\t' not in para_text:
                found, tc_text = find_tc_field_in_paragraph(para, 'List of Abbreviations')
                if found:
                    abbrev_tc_found = True
                    abbrev_tc_text = tc_text
                    break

        if abbrev_tc_found:
            print("PASS: Component 2 — TC field found for 'List of Abbreviations' at level 1 (0.4 pts)")
            print("  instrText: " + str(abbrev_tc_text))
            total_score += 0.4
        else:
            print("FAIL: Component 2 — No TC field at level 1 found in 'List of Abbreviations' paragraph")
    except Exception as e:
        print("ERROR: Component 2 — " + str(e))

    # -----------------------------------------------------------------------
    # Component 3: TOC section updated — both entries appear in TOC (0.2 pts)
    # The TOC is a manual list near the beginning of the document.
    # After update, both entries should appear as TOC lines (text + tab + page number).
    # This MUST fail on initial_env (TOC has only 6 entries) and pass on golden_env
    # -----------------------------------------------------------------------
    try:
        toc_glossary_found = False
        toc_abbrev_found = False

        # Scan paragraphs in the TOC region (near the start of the document)
        # The TOC section starts after the 'Table of Contents' heading
        # We look for lines with a tab character (indicating TOC entry format)
        in_toc = False
        for para in doc.paragraphs:
            para_text = para.text.strip()
            style_name = para.style.name

            # Enter TOC region when we see the heading
            if style_name == 'Heading 1' and 'Table of Contents' in para_text:
                in_toc = True
                continue

            # Exit TOC region when we hit the next Heading 1
            if in_toc and style_name == 'Heading 1':
                break

            if in_toc:
                # TOC entries have a tab character separating text from page number
                if '\t' in para_text:
                    entry_title = para_text.split('\t')[0].strip()
                    if entry_title == 'Glossary of Terms':
                        toc_glossary_found = True
                        print("  TOC entry found: " + repr(para_text))
                    elif entry_title == 'List of Abbreviations':
                        toc_abbrev_found = True
                        print("  TOC entry found: " + repr(para_text))

        if toc_glossary_found and toc_abbrev_found:
            print("PASS: Component 3 — TOC section contains both 'Glossary of Terms' and 'List of Abbreviations' entries (0.2 pts)")
            total_score += 0.2
        elif toc_glossary_found:
            print("FAIL: Component 3 — TOC contains 'Glossary of Terms' but missing 'List of Abbreviations'")
        elif toc_abbrev_found:
            print("FAIL: Component 3 — TOC contains 'List of Abbreviations' but missing 'Glossary of Terms'")
        else:
            print("FAIL: Component 3 — TOC section does not contain either entry")
    except Exception as e:
        print("ERROR: Component 3 — " + str(e))

    final_score = min(total_score, 1.0)
    print()
    print("Score: " + str(total_score) + "/1.0")
    print("REWARD: " + str(final_score))
    return final_score


# Default: test against canonical artifact path in the VM env
file_path = WORKDIR + '/' + TASK_ID + '.docx'
if not os.path.exists(file_path):
    print("File not found: " + file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
