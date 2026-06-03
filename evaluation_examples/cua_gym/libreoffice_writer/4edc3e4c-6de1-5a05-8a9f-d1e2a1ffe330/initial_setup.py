"""
Initial Setup: Insert custom TOC entries for non-heading paragraphs
Task ID: writer_struct_038
Domain: libreoffice_writer

Creates a 12-page technical standard document with:
- A TOC at the beginning listing 6 Heading 1 entries
- Body content with 6 sections (Heading 1 style)
- 'Glossary of Terms' and 'List of Abbreviations' on pages 10-11
  styled as Default Paragraph Style (NOT headings, NO TOC index marks)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_struct_038'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_toc_field(doc):
    """Insert a TOC field (Table of Contents) into the document."""
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = para.add_run()
    # Begin field
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    # instrText for TOC (Table of Contents, Heading levels 1-3)
    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    # Separate
    run3 = para.add_run()
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar_sep)

    # TOC result entries (displayed text - these will be the 6 heading entries)
    toc_entries = [
        ("1. Introduction", "1"),
        ("2. Scope and Purpose", "2"),
        ("3. Technical Requirements", "4"),
        ("4. Implementation Guidelines", "6"),
        ("5. Testing and Validation", "8"),
        ("6. Compliance and Certification", "10"),
    ]
    for entry_text, page_num in toc_entries:
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run(f"{entry_text}\t{page_num}")

    # End field
    run4 = para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run4._r.append(fldChar_end)

    return para


def add_section_content(doc, heading_text, level=1, paragraphs_count=3, page_break_before=False):
    """Add a heading and content paragraphs."""
    if page_break_before:
        doc.add_page_break()
    heading = doc.add_heading(heading_text, level=level)
    content_data = [
        "This section provides detailed information regarding the technical specifications and "
        "operational parameters required for compliance with the established standards.",
        "All implementations must adhere to the guidelines set forth in this document. "
        "Deviations from these specifications require formal written approval from the "
        "standards committee before implementation.",
        "The requirements outlined here supersede any previous documentation and should be "
        "treated as the authoritative reference for all related activities.",
        "Organizations implementing these standards should ensure that their internal processes "
        "align with the documented procedures and that all personnel are adequately trained.",
        "Regular audits and reviews are recommended to ensure continued compliance with the "
        "standards as they evolve over time.",
    ]
    for i in range(paragraphs_count):
        doc.add_paragraph(content_data[i % len(content_data)])


def create_initial():
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # --- Title Page ---
    title_para = doc.add_heading("Technical Standards Document", level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("ISO/IEC 27001:2022 — Information Security Management Systems")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    version = doc.add_paragraph("Version 3.2 | Revision Date: March 2025")
    version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    # --- Table of Contents heading ---
    toc_heading = doc.add_heading("Table of Contents", level=1)

    # TOC entries as plain text (simulating the TOC - no TC marks here)
    toc_lines = [
        ("1. Introduction", "1"),
        ("2. Scope and Purpose", "2"),
        ("3. Technical Requirements", "4"),
        ("4. Implementation Guidelines", "6"),
        ("5. Testing and Validation", "8"),
        ("6. Compliance and Certification", "10"),
    ]
    for entry, page in toc_lines:
        toc_entry = doc.add_paragraph(style='Normal')
        run = toc_entry.add_run(f"{entry}")
        run.font.size = Pt(11)
        tab_run = toc_entry.add_run(f"\t{page}")
        tab_run.font.size = Pt(11)

    doc.add_page_break()

    # --- Section 1: Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document establishes the technical standards and requirements for information "
        "security management systems within the organization. These standards have been "
        "developed in accordance with international best practices and regulatory requirements."
    )
    doc.add_paragraph(
        "The purpose of this standard is to provide a comprehensive framework for managing "
        "information security risks and ensuring the confidentiality, integrity, and availability "
        "of organizational information assets."
    )
    doc.add_paragraph(
        "This standard applies to all organizational units, subsidiaries, and contractors who "
        "handle sensitive information on behalf of the organization. Compliance with these "
        "standards is mandatory for all covered entities."
    )

    # --- Section 2: Scope and Purpose ---
    doc.add_heading("2. Scope and Purpose", level=1)
    doc.add_paragraph(
        "The scope of this standard encompasses all information systems, networks, and "
        "data repositories operated by or on behalf of the organization. This includes "
        "cloud-based services, on-premises infrastructure, and hybrid environments."
    )
    doc.add_paragraph(
        "The primary purpose is to establish minimum security requirements that must be "
        "met by all systems processing, storing, or transmitting organizational data. "
        "These requirements are designed to mitigate identified risks and protect against "
        "known threat vectors."
    )
    doc.add_paragraph(
        "Secondary objectives include facilitating regulatory compliance, enabling consistent "
        "security assessments, and providing a foundation for continuous improvement of "
        "the organization's security posture."
    )

    doc.add_page_break()

    # --- Section 3: Technical Requirements ---
    doc.add_heading("3. Technical Requirements", level=1)
    doc.add_paragraph(
        "All systems must implement multi-factor authentication for administrative access. "
        "Password policies must enforce minimum complexity requirements including a minimum "
        "length of 12 characters, uppercase and lowercase letters, numbers, and special characters."
    )
    doc.add_paragraph(
        "Network segmentation must be implemented to isolate critical systems from general "
        "network traffic. Firewalls and intrusion detection systems must be deployed at all "
        "network boundaries and configured according to the principle of least privilege."
    )
    doc.add_paragraph(
        "All data in transit must be encrypted using TLS 1.2 or higher. Data at rest "
        "must be encrypted using AES-256 or equivalent. Encryption keys must be managed "
        "through approved key management systems and rotated according to established schedules."
    )
    doc.add_paragraph(
        "Vulnerability management programs must include regular scanning (at minimum weekly "
        "for critical systems), penetration testing (at minimum annually), and a defined "
        "remediation timeline based on CVSS severity scores."
    )

    doc.add_page_break()

    # --- Section 4: Implementation Guidelines ---
    doc.add_heading("4. Implementation Guidelines", level=1)
    doc.add_paragraph(
        "Implementation of these standards should follow a risk-based approach, prioritizing "
        "controls for systems handling the most sensitive information or those exposed to the "
        "greatest threat landscape. A formal risk assessment must precede implementation planning."
    )
    doc.add_paragraph(
        "Project teams responsible for implementation must include qualified security personnel "
        "with relevant certifications (CISSP, CISM, or equivalent). External consultants may "
        "supplement internal expertise but cannot replace the requirement for internal oversight."
    )
    doc.add_paragraph(
        "Implementation timelines must be documented and approved by the Chief Information "
        "Security Officer (CISO). Milestone reviews are required at 25%, 50%, 75%, and "
        "100% completion points. Any deviations from the approved timeline must be formally reported."
    )
    doc.add_paragraph(
        "Change management procedures must be followed for all implementations. Emergency "
        "changes require post-implementation review within 72 hours. All changes must be "
        "documented in the organization's configuration management database (CMDB)."
    )

    doc.add_page_break()

    # --- Section 5: Testing and Validation ---
    doc.add_heading("5. Testing and Validation", level=1)
    doc.add_paragraph(
        "All security controls must be tested prior to deployment in production environments. "
        "Testing must include both functional testing to verify controls operate as intended "
        "and negative testing to confirm that unauthorized actions are properly blocked."
    )
    doc.add_paragraph(
        "User acceptance testing (UAT) must include security test cases developed by the "
        "security team. Sign-off from both the business owner and the security team is required "
        "before any system can be moved to production status."
    )
    doc.add_paragraph(
        "Post-implementation testing must be conducted within 30 days of deployment. This "
        "testing should include review of security logs, verification of monitoring alerts, "
        "and confirmation that all planned controls are operational."
    )

    doc.add_page_break()

    # --- Section 6: Compliance and Certification ---
    doc.add_heading("6. Compliance and Certification", level=1)
    doc.add_paragraph(
        "Annual compliance assessments are required for all systems in scope. Assessments "
        "must be conducted by qualified assessors who are independent of the system owners. "
        "Assessment reports must be reviewed by the CISO and presented to executive leadership."
    )
    doc.add_paragraph(
        "Systems that fail compliance assessments must implement a remediation plan within "
        "30 days. Continued non-compliance after 90 days may result in system decommissioning "
        "or escalation to the board of directors."
    )
    doc.add_paragraph(
        "Third-party certification against ISO 27001 is strongly recommended for critical "
        "systems and required for systems processing customer payment data. Certification "
        "must be maintained through regular surveillance audits and recertification cycles."
    )

    doc.add_page_break()

    # --- Appendix pages (to reach page 10) ---
    doc.add_heading("Appendix A: Security Control Reference", level=1)
    doc.add_paragraph(
        "This appendix provides reference information for the security controls referenced "
        "throughout this document. Controls are mapped to relevant industry frameworks "
        "including NIST CSF, ISO 27001, and SOC 2 Type II."
    )
    doc.add_paragraph(
        "Control ID: SEC-001 — Access Control\n"
        "Description: All systems must implement role-based access control (RBAC) with "
        "the principle of least privilege. Access rights must be reviewed quarterly.\n"
        "Framework Mapping: ISO 27001 A.9, NIST CSF PR.AC-1, SOC 2 CC6.1"
    )
    doc.add_paragraph(
        "Control ID: SEC-002 — Cryptographic Controls\n"
        "Description: Approved cryptographic algorithms must be used for data protection. "
        "See Appendix B for the approved algorithm list.\n"
        "Framework Mapping: ISO 27001 A.10, NIST CSF PR.DS-2, SOC 2 CC6.7"
    )
    doc.add_paragraph(
        "Control ID: SEC-003 — Incident Response\n"
        "Description: Documented incident response procedures must be maintained and tested "
        "through tabletop exercises at least annually.\n"
        "Framework Mapping: ISO 27001 A.16, NIST CSF RS.RP-1, SOC 2 CC7.3"
    )

    doc.add_page_break()

    doc.add_heading("Appendix B: Approved Cryptographic Algorithms", level=1)
    doc.add_paragraph(
        "The following cryptographic algorithms have been approved for use within the "
        "organization's information systems. Use of non-approved algorithms requires "
        "formal exception approval from the CISO."
    )
    # Add a table for algorithms
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Algorithm'
    hdr_cells[1].text = 'Use Case'
    hdr_cells[2].text = 'Key Size'
    for r in hdr_cells:
        for para in r.paragraphs:
            for run in para.runs:
                run.font.bold = True

    algo_data = [
        ('AES', 'Symmetric encryption', '128, 192, or 256 bits'),
        ('RSA', 'Asymmetric encryption / key exchange', '2048 bits minimum'),
        ('ECDSA', 'Digital signatures', 'P-256 or P-384 curve'),
        ('SHA-256', 'Hashing', 'N/A'),
        ('SHA-3', 'Hashing (alternative)', 'N/A'),
        ('TLS 1.2/1.3', 'Transport security', 'Per cipher suite'),
    ]
    for algo, use_case, key_size in algo_data:
        row = table.add_row()
        row.cells[0].text = algo
        row.cells[1].text = use_case
        row.cells[2].text = key_size

    doc.add_paragraph()

    doc.add_page_break()

    # --- Glossary of Terms (page 10) - styled as Default Paragraph Style, NOT heading ---
    # This paragraph should NOT have TOC index marks in the initial state
    glossary_para = doc.add_paragraph('Glossary of Terms')
    glossary_para.style = doc.styles['Normal']
    glossary_para.runs[0].font.size = Pt(16)
    glossary_para.runs[0].font.bold = True

    doc.add_paragraph(
        "The following terms are used throughout this document. Understanding these "
        "definitions is essential for correct interpretation of the requirements."
    )

    glossary_items = [
        ("Asset", "Any item of value to the organization, including information, software, hardware, and physical items."),
        ("Confidentiality", "The property that information is not made available or disclosed to unauthorized individuals, entities, or processes."),
        ("Control", "A measure that modifies risk, including policies, procedures, guidelines, practices, or organizational structures."),
        ("Integrity", "The property of accuracy and completeness of assets."),
        ("Risk", "The potential that a given threat will exploit vulnerabilities of an asset or group of assets."),
        ("Threat", "A potential cause of an unwanted incident that may result in harm to a system or organization."),
        ("Vulnerability", "A weakness of an asset or control that can be exploited by one or more threats."),
    ]
    for term, definition in glossary_items:
        p = doc.add_paragraph()
        term_run = p.add_run(f"{term}: ")
        term_run.font.bold = True
        def_run = p.add_run(definition)

    doc.add_page_break()

    # --- List of Abbreviations (page 11) - styled as Default Paragraph Style, NOT heading ---
    # This paragraph should NOT have TOC index marks in the initial state
    abbrev_para = doc.add_paragraph('List of Abbreviations')
    abbrev_para.style = doc.styles['Normal']
    abbrev_para.runs[0].font.size = Pt(16)
    abbrev_para.runs[0].font.bold = True

    doc.add_paragraph(
        "The following abbreviations appear throughout this document:"
    )

    abbreviations = [
        ("AES", "Advanced Encryption Standard"),
        ("CISO", "Chief Information Security Officer"),
        ("CISSP", "Certified Information Systems Security Professional"),
        ("CMDB", "Configuration Management Database"),
        ("CVSS", "Common Vulnerability Scoring System"),
        ("ECDSA", "Elliptic Curve Digital Signature Algorithm"),
        ("IDS", "Intrusion Detection System"),
        ("ISO", "International Organization for Standardization"),
        ("MFA", "Multi-Factor Authentication"),
        ("NIST", "National Institute of Standards and Technology"),
        ("RBAC", "Role-Based Access Control"),
        ("RSA", "Rivest–Shamir–Adleman"),
        ("SHA", "Secure Hash Algorithm"),
        ("SOC", "Service Organization Control"),
        ("TLS", "Transport Layer Security"),
        ("UAT", "User Acceptance Testing"),
    ]
    for abbr, full_form in abbreviations:
        p = doc.add_paragraph()
        abbr_run = p.add_run(f"{abbr}: ")
        abbr_run.font.bold = True
        full_run = p.add_run(full_form)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup - open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
