"""
Initial Setup: Product price list with default tab stops (no decimal alignment)
Task ID: wrpara_024
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'wrpara_024'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title heading
    heading = doc.add_heading("Product Price List", level=1)

    # Intro paragraph
    intro = doc.add_paragraph(
        "Below is our current product catalog with retail pricing. "
        "All prices are listed in USD and include applicable taxes."
    )
    intro.paragraph_format.space_after = Pt(12)

    # Product lines - each with product name, tab, then price
    # Using default tab stops only (no decimal alignment)
    products = [
        ("Premium Widget A", "$12.50"),
        ("Economy Bolt Pack", "$9.99"),
        ("Industrial Gear Assembly", "$125.00"),
        ("Micro Fastener Kit", "$3.75"),
        ("Standard Bracket Set", "$42.10"),
        ("Heavy-Duty Motor Unit", "$1,250.00"),
    ]

    for name, price in products:
        para = doc.add_paragraph()
        run_name = para.add_run(name)
        run_name.font.name = "Calibri"
        run_name.font.size = Pt(11)
        run_tab = para.add_run("\t")
        run_price = para.add_run(price)
        run_price.font.name = "Calibri"
        run_price.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(2)

    # Footer note
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    run_note = note.add_run("Prices effective as of March 2025. Subject to change without notice.")
    run_note.font.italic = True
    run_note.font.size = Pt(9)
    run_note.font.name = "Calibri"

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the file
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
