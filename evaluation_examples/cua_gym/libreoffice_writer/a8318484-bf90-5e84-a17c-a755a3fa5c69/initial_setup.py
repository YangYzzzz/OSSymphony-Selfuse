"""
Initial Setup: Workplace safety notice document with no page border.
Task ID: writer_page_063
Domain: libreoffice_writer

Creates: /home/user/Desktop/safety_notice.docx
- 1-page workplace safety notice
- Page: A4, portrait, margins top/bottom/left/right = 2.54cm
- No page border (agent must add it)
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'
TASK_ID = 'safety_notice'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Page Setup: A4, portrait, margins 2.54cm ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4 width
    section.page_height = Cm(29.7)   # A4 height
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # --- Ensure no page border in sectPr ---
    # Remove any existing pgBorders element if present
    sect_pr = section._sectPr
    for pb in sect_pr.findall(qn('w:pgBorders')):
        sect_pr.remove(pb)

    # --- Document content: 1-page workplace safety notice ---

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run('WORKPLACE SAFETY NOTICE')
    title_run.bold = True
    title_run.font.size = Pt(18)

    doc.add_paragraph()  # spacing

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = sub_para.add_run('All Staff — Please Read and Comply')
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    doc.add_paragraph()  # spacing

    # Introduction
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'This notice outlines mandatory safety procedures applicable to all employees, '
        'contractors, and visitors operating on company premises. Compliance with these '
        'guidelines is required at all times.'
    )
    intro_run.font.size = Pt(11)

    doc.add_paragraph()

    # Section: Personal Protective Equipment
    ppe_heading = doc.add_paragraph()
    ppe_h_run = ppe_heading.add_run('1. Personal Protective Equipment (PPE)')
    ppe_h_run.bold = True
    ppe_h_run.font.size = Pt(12)

    ppe_body = doc.add_paragraph()
    ppe_body.add_run(
        'All personnel entering the warehouse or production floor must wear appropriate PPE, '
        'including hard hats, high-visibility vests, steel-toe boots, and safety goggles. '
        'PPE must be inspected before each shift. Damaged equipment must be reported immediately '
        'to the Safety Officer (ext. 204).'
    ).font.size = Pt(11)

    doc.add_paragraph()

    # Section: Emergency Procedures
    emer_heading = doc.add_paragraph()
    emer_h_run = emer_heading.add_run('2. Emergency Procedures')
    emer_h_run.bold = True
    emer_h_run.font.size = Pt(12)

    emer_body = doc.add_paragraph()
    emer_body.add_run(
        'In the event of a fire, chemical spill, or other emergency:\n'
        '  \u2022 Activate the nearest fire alarm pull station.\n'
        '  \u2022 Evacuate immediately using the designated emergency exits.\n'
        '  \u2022 Proceed to the assembly point in Car Park B.\n'
        '  \u2022 Do not re-enter the building until the All Clear is given by the Fire Warden.'
    ).font.size = Pt(11)

    doc.add_paragraph()

    # Section: Incident Reporting
    inc_heading = doc.add_paragraph()
    inc_h_run = inc_heading.add_run('3. Incident Reporting')
    inc_h_run.bold = True
    inc_h_run.font.size = Pt(12)

    inc_body = doc.add_paragraph()
    inc_body.add_run(
        'Any injury, near-miss, or unsafe condition must be reported within 2 hours of occurrence. '
        'Complete Form HS-07 available from the HR portal and submit to your line manager and the '
        'Health & Safety team. Failure to report incidents may result in disciplinary action.'
    ).font.size = Pt(11)

    doc.add_paragraph()

    # Footer note
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_para.add_run(
        'Issued by: Health & Safety Department  |  Effective Date: 01 March 2025  |  Review Date: 01 March 2026'
    )
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
