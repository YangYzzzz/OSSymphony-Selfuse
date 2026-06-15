"""
Initial Setup: Generate a 40-page technical manual with properly formatted headings.
Task ID: writer_rd_029
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


# Technical manual content structure:
# 5 Heading 1 chapters, each with ~3 Heading 2 sections, some with Heading 3 subsections
# Total: 5 H1, 15 H2, 10 H3

MANUAL_STRUCTURE = [
    {
        "h1": "Chapter 1: System Architecture Overview",
        "sections": [
            {
                "h2": "1.1 Platform Components and Dependencies",
                "body": [
                    "The Meridian Analytics Platform is built on a microservices architecture that separates concerns across multiple independent services. Each service communicates through a well-defined REST API layer, ensuring loose coupling and high cohesion between components.",
                    "At the core of the system sits the Data Ingestion Service (DIS), which handles all incoming data streams from external providers. The DIS processes approximately 2.4 million records per hour during peak load, utilizing a combination of Apache Kafka message queues and Redis caching layers to maintain throughput. Connection pooling is managed through PgBouncer, with a maximum of 200 concurrent connections to the PostgreSQL cluster.",
                    "The Authentication and Authorization Module (AAM) implements OAuth 2.0 with PKCE flow for all client applications. Token rotation occurs every 3600 seconds, with refresh tokens valid for 30 days. The module integrates with both Active Directory and LDAP for enterprise single sign-on capabilities.",
                ],
                "subsections": [
                    {
                        "h3": "1.1.1 Network Topology",
                        "body": [
                            "The production network is segmented into three primary zones: the DMZ (demilitarized zone) hosting the API gateway and load balancers, the application tier running containerized services on Kubernetes, and the data tier containing the database clusters and object storage. All inter-zone communication is encrypted using TLS 1.3 with mutual authentication.",
                            "Load balancing is performed by HAProxy instances configured in active-passive mode. Health checks run every 5 seconds, with a failure threshold of 3 consecutive missed checks before a node is removed from the rotation. The failover time between primary and secondary load balancers is under 2 seconds.",
                        ],
                    },
                    {
                        "h3": "1.1.2 Service Discovery and Registration",
                        "body": [
                            "Consul is used for service discovery across all environments. Each microservice registers itself on startup and sends heartbeat signals every 10 seconds. If a service fails to respond within 30 seconds, Consul automatically deregisters it and triggers an alert through the monitoring pipeline.",
                            "Service mesh communication is handled by Envoy sidecar proxies, which provide automatic retries, circuit breaking, and distributed tracing. The circuit breaker trips after 5 consecutive failures and remains open for 30 seconds before transitioning to half-open state.",
                        ],
                    },
                ],
            },
            {
                "h2": "1.2 Data Flow and Processing Pipeline",
                "body": [
                    "Data enters the platform through multiple ingestion points: REST APIs for batch uploads, WebSocket connections for real-time streaming, and SFTP endpoints for legacy system integration. Each ingestion point validates incoming data against predefined JSON schemas before forwarding to the processing pipeline.",
                    "The processing pipeline consists of four stages: validation, enrichment, transformation, and loading. Each stage is implemented as an independent microservice that can be scaled horizontally. The pipeline processes data in micro-batches of 1000 records, achieving an end-to-end latency of under 500 milliseconds for 95th percentile requests.",
                    "Error handling follows a dead letter queue pattern. Records that fail validation are routed to a separate Kafka topic for manual review. The system retains failed records for 90 days, after which they are archived to cold storage in AWS S3 Glacier.",
                ],
                "subsections": [],
            },
            {
                "h2": "1.3 Security Architecture and Compliance",
                "body": [
                    "The platform implements defense-in-depth security with multiple layers of protection. Network-level security is enforced through AWS Security Groups and Network ACLs. Application-level security includes input validation, output encoding, and parameterized queries to prevent injection attacks.",
                    "All data at rest is encrypted using AES-256 with keys managed through AWS Key Management Service (KMS). Customer-managed keys (CMK) are rotated annually, and the encryption context includes the tenant identifier to prevent cross-tenant data access. Data in transit uses TLS 1.3 exclusively; older protocol versions have been disabled across all endpoints.",
                    "Compliance with SOC 2 Type II, HIPAA, and GDPR requirements is maintained through automated audit logging. Every API call generates an audit event containing the user identity, action performed, resource accessed, timestamp, and source IP address. Audit logs are stored in an immutable append-only database with a retention period of 7 years.",
                ],
                "subsections": [
                    {
                        "h3": "1.3.1 Access Control Policies",
                        "body": [
                            "Role-based access control (RBAC) is implemented with five predefined roles: Viewer, Analyst, Editor, Admin, and Super Admin. Each role has a specific set of permissions mapped to API endpoints. Custom roles can be created by combining individual permissions, but the total number of custom roles per organization is limited to 50.",
                            "Multi-factor authentication (MFA) is mandatory for all users with Admin or Super Admin privileges. The system supports TOTP-based authenticators, hardware security keys (FIDO2/WebAuthn), and SMS verification as a fallback option. Session tokens are invalidated after 8 hours of inactivity.",
                        ],
                    },
                ],
            },
        ],
    },
    {
        "h1": "Chapter 2: Database Administration Guide",
        "sections": [
            {
                "h2": "2.1 PostgreSQL Cluster Configuration",
                "body": [
                    "The primary database cluster runs PostgreSQL 15.4 with streaming replication to two synchronous standby nodes. The cluster is deployed on dedicated EC2 instances (r6g.4xlarge) with 128 GB RAM and 2 TB NVMe SSD storage. Connection pooling through PgBouncer maintains a pool size of 100 connections per application service.",
                    "Database partitioning is implemented for all time-series tables using declarative range partitioning on the timestamp column. Partitions are created monthly and retained for 24 months before being migrated to the archive schema. The partition management job runs daily at 02:00 UTC and creates partitions 3 months in advance.",
                    "Vacuum operations are configured with aggressive settings to prevent table bloat. Autovacuum triggers when dead tuples exceed 10% of live tuples or 50,000 dead tuples, whichever comes first. The vacuum cost delay is set to 2 milliseconds with a cost limit of 2000, balancing cleanup performance against query workload impact.",
                ],
                "subsections": [
                    {
                        "h3": "2.1.1 Replication and Failover Procedures",
                        "body": [
                            "Patroni manages the PostgreSQL high availability cluster with etcd as the distributed configuration store. Failover is automatic when the primary node becomes unreachable for more than 30 seconds. The maximum replication lag threshold for synchronous standbys is configured at 1 MB to prevent data loss during failover.",
                            "Manual failover can be initiated through the Patroni REST API or the patronictl command-line tool. Before performing a planned failover, operators should verify that all standbys are caught up by checking the pg_stat_replication view. The expected switchover time is under 10 seconds for planned failovers.",
                        ],
                    },
                ],
            },
            {
                "h2": "2.2 Backup and Recovery Strategies",
                "body": [
                    "Full base backups are taken weekly using pgBackRest with parallel compression (zstd level 3). Incremental backups run daily at 03:00 UTC, and WAL archiving ensures point-in-time recovery capability. Backup retention is configured for 4 weeks of full backups and 90 days of WAL archives.",
                    "Recovery testing is performed monthly in an isolated environment. The recovery procedure involves restoring the most recent full backup, applying incremental backups, and replaying WAL segments to the target recovery point. The mean time to recovery (MTTR) target is 30 minutes for the complete database cluster.",
                    "Cross-region backup replication ensures disaster recovery capability. Backups are replicated to a secondary AWS region (eu-west-1) with a maximum replication delay of 1 hour. The cross-region recovery procedure has been tested quarterly with a target Recovery Point Objective (RPO) of 1 hour and Recovery Time Objective (RTO) of 4 hours.",
                ],
                "subsections": [],
            },
            {
                "h2": "2.3 Performance Tuning and Monitoring",
                "body": [
                    "Query performance is monitored through pg_stat_statements, which tracks execution statistics for all normalized queries. The top 50 queries by total execution time are reviewed weekly, and execution plans are analyzed for any query with a mean execution time exceeding 100 milliseconds.",
                    "Index maintenance includes regular REINDEX operations for bloated indexes and CREATE INDEX CONCURRENTLY for new index additions. The index advisor runs monthly to suggest new indexes based on query patterns and to identify unused indexes that can be safely dropped.",
                    "Connection monitoring alerts are configured for connection pool saturation (above 80%), long-running transactions (exceeding 5 minutes), and lock wait times (exceeding 10 seconds). The monitoring dashboard provides real-time visibility into query throughput, latency percentiles, and resource utilization.",
                ],
                "subsections": [
                    {
                        "h3": "2.3.1 Query Optimization Best Practices",
                        "body": [
                            "All new queries must be accompanied by an execution plan showing the expected cost and row estimates. Queries that perform sequential scans on tables larger than 10,000 rows must include a justification or an appropriate index. The use of SELECT * is prohibited in production code; only required columns should be specified.",
                            "Common table expressions (CTEs) should be used judiciously, as PostgreSQL may materialize them, preventing predicate pushdown. For large result sets, consider using lateral joins or subqueries instead. The query planner's work_mem setting can be adjusted per-session for complex analytical queries that require hash joins or sorts.",
                        ],
                    },
                ],
            },
        ],
    },
    {
        "h1": "Chapter 3: API Development Standards",
        "sections": [
            {
                "h2": "3.1 RESTful Endpoint Design Conventions",
                "body": [
                    "All API endpoints follow REST conventions with resource-based URLs and standard HTTP methods. Resource names use plural nouns in kebab-case (e.g., /api/v2/data-sources). Nested resources are limited to one level of depth to maintain URL readability.",
                    "API versioning follows the URL path strategy with major version numbers only (v1, v2). Breaking changes require a new major version, while backward-compatible additions are published under the existing version. Deprecated endpoints include a Sunset header indicating the removal date, with a minimum deprecation period of 6 months.",
                    "Pagination is implemented using cursor-based pagination for all list endpoints. The response includes a next_cursor field and supports a page_size parameter with a maximum of 100 items per page. The total count is provided in a separate X-Total-Count header to avoid expensive COUNT queries on large tables.",
                ],
                "subsections": [],
            },
            {
                "h2": "3.2 Request Validation and Error Handling",
                "body": [
                    "Input validation is performed at three levels: schema validation using JSON Schema Draft 2020-12, business rule validation in the service layer, and database constraint validation at the persistence layer. All validation errors return HTTP 422 with a standardized error response body.",
                    "The error response format includes a machine-readable error code, human-readable message, field-level details array, and a request correlation ID. Error codes follow a hierarchical naming convention (e.g., VALIDATION.FIELD.REQUIRED, AUTH.TOKEN.EXPIRED) to facilitate programmatic error handling by API consumers.",
                    "Rate limiting is enforced at the API gateway level using a token bucket algorithm. Default limits are 1000 requests per minute for authenticated users and 100 requests per minute for anonymous access. Rate limit headers (X-RateLimit-Limit, X-RateLimit-Remaining, X-RateLimit-Reset) are included in all responses.",
                ],
                "subsections": [
                    {
                        "h3": "3.2.1 Circuit Breaker Implementation",
                        "body": [
                            "External service calls are protected by circuit breakers implemented using the Resilience4j library. The circuit breaker monitors the failure rate over a sliding window of 100 requests. When the failure rate exceeds 50%, the circuit opens and all subsequent calls fail fast for 60 seconds.",
                            "During the half-open state, 10 probe requests are allowed through. If the success rate of these probes exceeds 80%, the circuit closes; otherwise, it returns to the open state. Fallback responses are configured for each external dependency to ensure graceful degradation.",
                        ],
                    },
                ],
            },
            {
                "h2": "3.3 Authentication and Token Management",
                "body": [
                    "JWT tokens are issued by the authorization server with RS256 signing using a 2048-bit RSA key pair. Access tokens have a lifetime of 3600 seconds and contain the user ID, tenant ID, roles, and a session fingerprint. The token payload is kept minimal to reduce bandwidth overhead on every API call.",
                    "Refresh tokens are opaque strings stored in a Redis-backed session store. They support token rotation: each refresh request invalidates the previous refresh token and issues a new pair. If a previously used refresh token is presented, all tokens for that session are immediately revoked as a security precaution.",
                    "Service-to-service authentication uses mTLS with client certificates issued by an internal Certificate Authority. Certificate rotation is automated through cert-manager in Kubernetes, with certificates valid for 90 days and renewal triggered 30 days before expiration.",
                ],
                "subsections": [],
            },
        ],
    },
    {
        "h1": "Chapter 4: Deployment and Operations",
        "sections": [
            {
                "h2": "4.1 Container Orchestration with Kubernetes",
                "body": [
                    "The production Kubernetes cluster runs on Amazon EKS with managed node groups. The cluster consists of three node pools: general-purpose (m6i.2xlarge), memory-optimized (r6i.4xlarge), and GPU-enabled (g5.xlarge) for machine learning workloads. Cluster autoscaler adjusts node counts based on pending pod resource requests.",
                    "Namespace isolation is enforced through Network Policies and Resource Quotas. Each team operates within a dedicated namespace with CPU and memory limits. Inter-namespace communication is denied by default and must be explicitly allowed through NetworkPolicy objects.",
                    "Helm charts are maintained in a centralized chart repository and follow semantic versioning. Each service has its own chart with configurable values for replicas, resource limits, environment variables, and feature flags. Chart updates are tested in a staging environment before promotion to production.",
                ],
                "subsections": [
                    {
                        "h3": "4.1.1 Pod Resource Management",
                        "body": [
                            "Resource requests and limits are mandatory for all containers. CPU requests are set to the P95 utilization observed during load testing, and memory limits are set to 150% of the maximum observed usage. The Vertical Pod Autoscaler (VPA) provides recommendations but operates in recommendation-only mode to prevent unexpected restarts.",
                            "Horizontal Pod Autoscaler (HPA) is configured for all stateless services with a target CPU utilization of 70%. The scaling behavior includes a stabilization window of 300 seconds for scale-down to prevent thrashing during transient load spikes. Custom metrics from Prometheus can be used for scaling decisions when CPU-based scaling is insufficient.",
                        ],
                    },
                ],
            },
            {
                "h2": "4.2 CI/CD Pipeline Configuration",
                "body": [
                    "The continuous integration pipeline is implemented in GitHub Actions with matrix builds for multiple target environments. Each pull request triggers unit tests, integration tests, static code analysis (SonarQube), dependency vulnerability scanning (Snyk), and container image scanning (Trivy).",
                    "Deployment to staging is automatic on merge to the develop branch. Production deployments require manual approval from two designated reviewers and are executed during the maintenance window (Tuesday and Thursday, 02:00-06:00 UTC). Blue-green deployment strategy is used for zero-downtime releases.",
                    "Rollback procedures are automated through ArgoCD with GitOps workflow. If the post-deployment health check fails within 10 minutes, ArgoCD automatically reverts to the previous stable revision. Manual rollback can be triggered through the ArgoCD dashboard or CLI with immediate effect.",
                ],
                "subsections": [
                    {
                        "h3": "4.2.1 Environment Promotion Strategy",
                        "body": [
                            "The promotion path follows four stages: development, staging, pre-production, and production. Each stage has specific quality gates that must be passed before promotion. The development environment accepts any build, staging requires all unit tests passing, pre-production requires integration and performance tests, and production requires security review sign-off.",
                            "Feature flags managed through LaunchDarkly control the rollout of new features independently of deployments. Progressive rollout starts at 1% of traffic, increases to 10% after 24 hours if no anomalies are detected, then to 50% and finally 100% over the course of a week. Automated rollback triggers if error rates exceed baseline by more than 2 standard deviations.",
                        ],
                    },
                ],
            },
            {
                "h2": "4.3 Monitoring and Alerting Framework",
                "body": [
                    "Observability is built on the three pillars: metrics (Prometheus/Grafana), logs (Elasticsearch/Kibana), and traces (Jaeger). All services export standard metrics through the Prometheus client library, including request duration histograms, error counters, and saturation gauges.",
                    "Alerting rules are defined in Prometheus Alertmanager with three severity levels: warning (notification to Slack), critical (page to on-call engineer), and emergency (escalation to engineering leadership). Alert fatigue is managed through aggregation rules, silence periods, and weekly alert review meetings.",
                    "Distributed tracing captures the full request lifecycle across all microservices. Trace sampling is set at 10% for normal traffic and 100% for requests exceeding the P99 latency threshold. The trace context is propagated through W3C TraceContext headers and includes custom baggage items for business context correlation.",
                ],
                "subsections": [],
            },
        ],
    },
    {
        "h1": "Chapter 5: Disaster Recovery and Business Continuity",
        "sections": [
            {
                "h2": "5.1 Recovery Point and Time Objectives",
                "body": [
                    "The platform maintains differentiated recovery objectives based on data classification. Tier 1 (critical financial data): RPO of 0 seconds with synchronous replication and RTO of 15 minutes. Tier 2 (operational data): RPO of 1 hour with asynchronous replication and RTO of 4 hours. Tier 3 (analytical data): RPO of 24 hours with daily backups and RTO of 24 hours.",
                    "Recovery procedures are documented in runbooks stored in Confluence and version-controlled in Git. Each runbook includes prerequisite checks, step-by-step instructions, verification procedures, and rollback steps. Runbooks are reviewed quarterly and updated after any incident that reveals gaps or inaccuracies.",
                    "Disaster recovery testing is conducted semi-annually through tabletop exercises and annual full failover drills. The most recent full failover drill achieved an actual RTO of 22 minutes for Tier 1 systems and 3.5 hours for Tier 2 systems, both within the defined objectives.",
                ],
                "subsections": [],
            },
            {
                "h2": "5.2 Multi-Region Failover Architecture",
                "body": [
                    "The active-passive multi-region setup designates us-east-1 as the primary region and eu-west-1 as the secondary region. Database replication uses PostgreSQL logical replication for real-time data synchronization. Static assets and configuration are replicated through S3 Cross-Region Replication with versioning enabled.",
                    "DNS failover is managed through Route 53 health checks with a TTL of 60 seconds. When the primary region fails health checks for three consecutive intervals, traffic is automatically routed to the secondary region. The DNS propagation delay means actual failover takes approximately 2-3 minutes for most clients.",
                    "Application state that cannot be replicated through the database (such as in-flight transactions and session caches) is handled through the compensating transaction pattern. After failover, the system enters a reconciliation phase that identifies and resolves any data inconsistencies between the regions.",
                ],
                "subsections": [
                    {
                        "h3": "5.2.1 Data Consistency Verification",
                        "body": [
                            "Post-failover data consistency is verified through a three-phase process. Phase 1: Row count comparison across all critical tables. Phase 2: Checksum verification of the last 24 hours of transactions. Phase 3: Business rule validation that checks aggregate metrics against known good values.",
                            "The consistency checker tool (pg_consistency_check) runs automatically after any failover event and generates a detailed report. Discrepancies exceeding the defined threshold (0.01% of total records) trigger a manual review by the database administration team. Historical consistency reports are retained for audit purposes.",
                        ],
                    },
                ],
            },
            {
                "h2": "5.3 Incident Response Procedures",
                "body": [
                    "The incident response process follows a five-phase model: detection, triage, containment, resolution, and post-mortem. Each phase has defined roles (Incident Commander, Technical Lead, Communications Lead) and time-boxed activities. Severity classification ranges from SEV-4 (minor degradation) to SEV-1 (complete service outage).",
                    "Communication during incidents follows a structured template posted to the #incident-response Slack channel. Status updates are provided every 15 minutes for SEV-1 and SEV-2 incidents, and every 30 minutes for SEV-3 incidents. External customer communication is handled through the status page (status.meridian-analytics.com) with automated updates.",
                    "Post-incident reviews are conducted within 5 business days of incident resolution. The review produces a blameless post-mortem document that includes a timeline of events, root cause analysis, impact assessment, and action items with assigned owners and due dates. Action items are tracked in Jira and reviewed in the weekly engineering leadership meeting.",
                ],
                "subsections": [],
            },
        ],
    },
]


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Heading styles configuration
    for level in [1, 2, 3]:
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Calibri'

    # Title page
    title_para = doc.add_heading('Meridian Analytics Platform', level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Technical Operations Manual')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    vrun = version_para.add_run('Version 3.2.1 — March 2025')
    vrun.font.size = Pt(12)
    vrun.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()
    doc.add_paragraph()

    confidential = doc.add_paragraph()
    confidential.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    crun = confidential.add_run('CONFIDENTIAL — Internal Use Only')
    crun.bold = True
    crun.font.size = Pt(10)

    # Page break after title page
    doc.add_page_break()

    # Build the manual content
    for chapter in MANUAL_STRUCTURE:
        doc.add_heading(chapter["h1"], level=1)

        for section in chapter["sections"]:
            doc.add_heading(section["h2"], level=2)

            for body_text in section["body"]:
                doc.add_paragraph(body_text)

            for subsection in section.get("subsections", []):
                doc.add_heading(subsection["h3"], level=3)

                for body_text in subsection["body"]:
                    doc.add_paragraph(body_text)

    # Add additional filler paragraphs to reach ~40 pages
    # Each page holds roughly 30-35 lines of 11pt text
    # We need extra content to push it to 40 pages
    filler_sections = [
        ("Appendix A: Configuration Reference", [
            "This appendix provides a complete reference of all configuration parameters used across the Meridian Analytics Platform. Parameters are organized by service and include default values, valid ranges, and descriptions of their effects on system behavior.",
            "Database Connection Parameters: The connection string follows the standard PostgreSQL URI format: postgresql://user:password@host:port/database?sslmode=require. The maximum connection pool size should not exceed the PostgreSQL max_connections setting divided by the number of application instances.",
            "Cache Configuration: Redis cache connections use a dedicated connection pool with a maximum of 50 connections per service instance. The default TTL for cached objects is 3600 seconds, with specific overrides for session data (28800 seconds) and rate limiting counters (60 seconds). Cache eviction follows the allkeys-lru policy.",
            "Logging Configuration: Log levels are configurable per package using the structured logging framework. Production environments default to INFO level, with DEBUG enabled for specific packages during troubleshooting. Log output format is JSON with fields for timestamp, level, service, trace_id, span_id, and message.",
            "Feature Flag Configuration: Feature flags are evaluated on every request using the LaunchDarkly SDK. The SDK maintains a local cache of flag values with a polling interval of 30 seconds. In case of connectivity loss to LaunchDarkly, the SDK falls back to cached values for up to 5 minutes before using hardcoded defaults.",
        ]),
        ("Appendix B: Troubleshooting Guide", [
            "This section provides solutions for common operational issues encountered during platform operation. Each entry includes symptoms, diagnostic steps, root cause analysis, and resolution procedures.",
            "Issue: High Memory Utilization on Application Pods — Symptoms: Pod memory usage exceeds 85% of limits, OOMKilled events in pod logs. Diagnosis: Check heap dump analysis with jmap, review recent deployment changes. Common causes include memory leaks in connection pools, excessive caching, or large batch processing jobs. Resolution: Restart affected pods, apply connection pool fixes, adjust memory limits if workload has legitimately increased.",
            "Issue: Database Connection Timeouts — Symptoms: Application logs show 'connection timeout' errors, PgBouncer stats show high wait queue. Diagnosis: Check pg_stat_activity for long-running queries, verify PgBouncer pool sizes, examine network connectivity. Common causes include unoptimized queries holding connections, PgBouncer pool exhaustion, or network partition between application and database tiers. Resolution: Kill long-running queries, increase pool sizes temporarily, investigate and fix the offending queries.",
            "Issue: Kafka Consumer Lag — Symptoms: Consumer group lag increasing steadily, data freshness SLA violations. Diagnosis: Check consumer group describe output, review consumer logs for processing errors. Common causes include slow downstream dependencies, deserialization failures, or insufficient consumer instances. Resolution: Scale consumer group, fix processing errors, implement circuit breakers for slow dependencies.",
            "Issue: Certificate Expiration Warnings — Symptoms: TLS handshake failures, cert-manager logs showing renewal failures. Diagnosis: Check certificate expiry dates with openssl, review cert-manager logs. Common causes include DNS challenges failing for Let's Encrypt, ACME rate limits, or misconfigured certificate issuers. Resolution: Verify DNS configuration, check rate limit status, manually renew if needed.",
            "Issue: Elasticsearch Index Performance Degradation — Symptoms: Search queries taking longer than 2 seconds, index refresh time increasing. Diagnosis: Check cluster health, review index segment counts, examine shard distribution. Common causes include unbalanced shards, too many small segments, or insufficient heap memory. Resolution: Force merge segments, rebalance shards, increase JVM heap up to 50% of available RAM.",
        ]),
        ("Appendix C: Security Compliance Checklist", [
            "This checklist is used during quarterly security audits to verify compliance with SOC 2, HIPAA, and GDPR requirements. Each item must be verified and signed off by the security team.",
            "Access Control Verification: Confirm that all user accounts have been reviewed within the last 90 days. Verify that terminated employee accounts are disabled within 24 hours of separation. Check that privileged access requires multi-factor authentication. Ensure that service accounts use certificate-based authentication and have minimum necessary permissions.",
            "Data Protection Verification: Confirm that all data at rest is encrypted with AES-256. Verify that encryption keys have been rotated within the last 12 months. Check that backup encryption uses separate keys from production. Ensure that data classification labels are applied to all databases and storage buckets.",
            "Network Security Verification: Confirm that all external-facing endpoints use TLS 1.3. Verify that internal service communication uses mutual TLS. Check that network segmentation policies are enforced. Ensure that intrusion detection systems are active and alerts are being reviewed within 4 hours.",
            "Logging and Monitoring Verification: Confirm that audit logs are enabled for all critical systems. Verify that log retention meets the 7-year requirement. Check that log integrity is protected through immutable storage. Ensure that security event monitoring is active and escalation procedures are tested monthly.",
            "Incident Response Verification: Confirm that the incident response plan has been reviewed within the last 6 months. Verify that all team members have completed incident response training. Check that communication channels and escalation contacts are current. Ensure that post-incident review action items from the last quarter have been completed.",
        ]),
        ("Appendix D: Performance Benchmarks", [
            "This appendix documents the baseline performance benchmarks established during the most recent load testing cycle (Q1 2025). These benchmarks serve as reference points for capacity planning and performance regression detection.",
            "API Response Time Benchmarks: GET /api/v2/data-sources — P50: 45ms, P95: 120ms, P99: 350ms. POST /api/v2/records — P50: 85ms, P95: 250ms, P99: 600ms. GET /api/v2/analytics/reports — P50: 200ms, P95: 800ms, P99: 2100ms. The analytics endpoint has higher latency due to complex aggregation queries.",
            "Throughput Benchmarks: The system sustains 5,000 requests per second under normal load across all endpoints. Peak load testing achieved 12,000 requests per second before error rates exceeded the 1% threshold. The data ingestion pipeline processes 2.4 million records per hour with a maximum burst capacity of 4.0 million records per hour.",
            "Database Performance Benchmarks: Simple key-value lookups complete in under 1ms. Complex joins across 3-4 tables complete in 50-200ms depending on table sizes. Analytical queries with window functions and aggregations complete in 500ms-3s for datasets up to 100 million rows. Index-only scans achieve sub-millisecond response times for properly indexed queries.",
            "Storage and I/O Benchmarks: Sequential write throughput on NVMe storage averages 1.8 GB/s. Random read IOPS averages 250,000 with 4KB block size. PostgreSQL checkpoint completion time averages 45 seconds with checkpoint_completion_target set to 0.9. WAL write throughput peaks at 200 MB/s during bulk data loads.",
            "Network Benchmarks: Inter-AZ latency averages 1.2ms (P99: 3.5ms). Cross-region latency between us-east-1 and eu-west-1 averages 75ms (P99: 120ms). Internal service-to-service call overhead (including TLS and serialization) adds approximately 2ms to each hop. The API gateway adds approximately 5ms of overhead for authentication and routing.",
        ]),
    ]

    for appendix_title, paragraphs in filler_sections:
        doc.add_heading(appendix_title, level=1)
        for para_text in paragraphs:
            doc.add_paragraph(para_text)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
