"""
Initial Setup: Create a Writer document with project plan content as plain paragraphs.
Task ID: writer_biz_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading("Nexus Platform Modernization — Project Plan", level=1)

    # Introduction paragraph
    doc.add_paragraph(
        "This document outlines the phased approach for modernizing the Nexus "
        "Platform infrastructure. Each phase includes key deliverables, milestones, "
        "and responsible teams. The timeline spans Q2 through Q4 2026 with a total "
        "budget allocation of $2.4M across all departments."
    )

    doc.add_paragraph()  # blank separator

    # Section heading
    heading = doc.add_heading("Project Phases", level=2)

    # Plain paragraphs for the phase items and sub-items (NO list formatting)
    # These are just normal paragraphs that the agent must convert to a numbered list
    doc.add_paragraph("Planning")
    doc.add_paragraph("Requirements gathering")
    doc.add_paragraph("Resource allocation")
    doc.add_paragraph("Execution")
    doc.add_paragraph("Development")
    doc.add_paragraph("Testing")
    doc.add_paragraph("Review")

    doc.add_paragraph()  # blank separator

    # Additional context paragraphs
    doc.add_heading("Stakeholders", level=2)
    doc.add_paragraph(
        "The project is sponsored by VP of Engineering, Diana Reyes, with oversight "
        "from the Architecture Review Board. Day-to-day management is handled by "
        "Senior PM, Kevin Okafor, who reports weekly to the steering committee."
    )

    doc.add_paragraph(
        "Cross-functional teams from Engineering, QA, DevOps, and Product Management "
        "will be involved throughout all phases. External consultants from Meridian "
        "Solutions will support the requirements gathering and architecture review."
    )

    doc.add_heading("Timeline", level=2)

    # A simple table for timeline
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Phase", "Start Date", "End Date"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        for run in table.cell(0, i).paragraphs[0].runs:
            run.bold = True
    data = [
        ["Planning", "April 15, 2026", "May 30, 2026"],
        ["Execution", "June 2, 2026", "September 26, 2026"],
        ["Review", "October 1, 2026", "October 31, 2026"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
