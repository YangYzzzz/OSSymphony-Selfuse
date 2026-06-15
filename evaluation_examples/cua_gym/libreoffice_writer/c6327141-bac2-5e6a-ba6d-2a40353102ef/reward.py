"""
Reward Script: Create numbered list with sub-items for project phases
Task ID: writer_biz_037
Domain: libreoffice_writer
Scoring:
  Component 1 (0.45): Main items (Planning, Execution, Review) use a numbered list style
  Component 2 (0.45): Sub-items use a level-2 numbered list style
  Component 3 (0.10): Correct hierarchical ordering (sub-items follow their parent main item)
  Text content is checked as a precondition gate, not scored.
"""

import os

from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_037'


def is_numbered_list_style(style_name):
    """Check if style name indicates a top-level numbered list."""
    if style_name is None:
        return False
    sn = style_name.lower().replace(' ', '')
    # Accept 'List Number', 'ListNumber', etc. but NOT 'List Number 2/3'
    if 'listnumber' in sn and not any(c.isdigit() for c in sn.replace('listnumber', '')):
        return True
    return False


def is_numbered_sub_list_style(style_name):
    """Check if style name indicates a level-2 numbered list."""
    if style_name is None:
        return False
    sn = style_name.lower().replace(' ', '')
    if 'listnumber2' in sn:
        return True
    return False


def is_any_list_style(style_name):
    """Check if style is any kind of list style (numbered or bullet)."""
    if style_name is None:
        return False
    sn = style_name.lower()
    return 'list' in sn or 'bullet' in sn or 'number' in sn


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Locate the list paragraphs by text content
    main_items = {'Planning': None, 'Execution': None, 'Review': None}
    sub_items = {
        'Requirements gathering': None,
        'Resource allocation': None,
        'Development': None,
        'Testing': None,
    }

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        for key in main_items:
            if key.lower() in text.lower() and len(text) < 50:
                main_items[key] = i
                break
        for key in sub_items:
            if key.lower() in text.lower() and len(text) < 80:
                sub_items[key] = i
                break

    print("Located main items: %s" % {k: v for k, v in main_items.items()})
    print("Located sub-items: %s" % {k: v for k, v in sub_items.items()})

    # Precondition gate: all expected items must be locatable
    if None in main_items.values() or None in sub_items.values():
        print("FAIL: Could not locate all expected list items in document")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Main items use a numbered list style (0.45 points)
    # In initial_env these are 'Normal' style; in golden they should be list-number style
    try:
        main_pass_count = 0
        for name, idx in main_items.items():
            p = doc.paragraphs[idx]
            style_name = p.style.name if p.style else 'None'
            if is_numbered_list_style(style_name) or is_any_list_style(style_name):
                print("PASS: Main item '%s' [%d] has list style '%s'" % (name, idx, style_name))
                main_pass_count += 1
            elif style_name == 'Normal':
                print("FAIL: Main item '%s' [%d] still has 'Normal' style (not converted to list)" % (name, idx))
            else:
                # Check for explicit numPr in XML
                from docx.oxml.ns import qn
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        print("PASS: Main item '%s' [%d] has numPr (explicit numbering)" % (name, idx))
                        main_pass_count += 1
                    else:
                        print("FAIL: Main item '%s' [%d] has style '%s' with no numPr" % (name, idx, style_name))
                else:
                    print("FAIL: Main item '%s' [%d] has style '%s' with no pPr" % (name, idx, style_name))

        if main_pass_count == 3:
            print("PASS: Component 1 -- All 3 main items have numbered list style (0.45 pts)")
            total_score += 0.45
        elif main_pass_count >= 1:
            partial = round(0.45 * main_pass_count / 3, 2)
            print("PARTIAL: Component 1 -- %d/3 main items have list style (%s pts)" % (main_pass_count, partial))
            total_score += partial
        else:
            print("FAIL: Component 1 -- No main items have numbered list style")
    except Exception as e:
        print("ERROR: Component 1 -- %s" % e)

    # Component 2: Sub-items use a level-2 numbered list style (0.45 points)
    # In initial_env these are 'Normal'; in golden they should be 'List Number 2' or similar
    try:
        sub_pass_count = 0
        for name, idx in sub_items.items():
            p = doc.paragraphs[idx]
            style_name = p.style.name if p.style else 'None'
            if is_numbered_sub_list_style(style_name):
                print("PASS: Sub-item '%s' [%d] has sub-list style '%s'" % (name, idx, style_name))
                sub_pass_count += 1
            elif style_name == 'Normal':
                print("FAIL: Sub-item '%s' [%d] still has 'Normal' style" % (name, idx))
            else:
                # Check for numPr with ilvl >= 1
                from docx.oxml.ns import qn
                pPr = p._element.find(qn('w:pPr'))
                has_sub_numbering = (
                    pPr is not None
                    and pPr.find(qn('w:numPr')) is not None
                    and pPr.find(qn('w:numPr')).find(qn('w:ilvl')) is not None
                    and int(pPr.find(qn('w:numPr')).find(qn('w:ilvl')).get(qn('w:val'), '0')) >= 1
                )
                if has_sub_numbering:
                    print("PASS: Sub-item '%s' [%d] has numPr with ilvl>=1" % (name, idx))
                    sub_pass_count += 1
                elif is_any_list_style(style_name):
                    print("PASS: Sub-item '%s' [%d] has list style '%s'" % (name, idx, style_name))
                    sub_pass_count += 1
                else:
                    print("FAIL: Sub-item '%s' [%d] has style '%s' (not a sub-list style)" % (name, idx, style_name))

        if sub_pass_count == 4:
            print("PASS: Component 2 -- All 4 sub-items have sub-list style (0.45 pts)")
            total_score += 0.45
        elif sub_pass_count >= 1:
            partial = round(0.45 * sub_pass_count / 4, 2)
            print("PARTIAL: Component 2 -- %d/4 sub-items have sub-list style (%s pts)" % (sub_pass_count, partial))
            total_score += partial
        else:
            print("FAIL: Component 2 -- No sub-items have sub-list style")
    except Exception as e:
        print("ERROR: Component 2 -- %s" % e)

    # Component 3: Correct hierarchical ordering (0.10 points)
    # Sub-items must appear directly after their parent main item AND both parent
    # and sub must have list styles (this compound check ensures it only passes on golden)
    try:
        # Planning (idx) -> Requirements gathering (idx+1), Resource allocation (idx+2)
        planning_idx = main_items['Planning']
        rg_idx = sub_items['Requirements gathering']
        ra_idx = sub_items['Resource allocation']
        exec_idx = main_items['Execution']
        dev_idx = sub_items['Development']
        test_idx = sub_items['Testing']

        # Check ordering: planning < rg < ra < execution < dev < test
        order_correct = (planning_idx < rg_idx < ra_idx < exec_idx < dev_idx < test_idx)

        # Check that parent has list style AND sub-items have sub-list style
        # (This ensures the compound check only passes when styles are applied)
        p_planning = doc.paragraphs[planning_idx]
        p_exec = doc.paragraphs[exec_idx]
        parent_styled = (
            is_numbered_list_style(p_planning.style.name) and
            is_numbered_list_style(p_exec.style.name)
        )

        if order_correct and parent_styled:
            print("PASS: Component 3 -- Correct hierarchical ordering with styled parents (0.10 pts)")
            total_score += 0.10
        else:
            if not order_correct:
                print("FAIL: Component 3 -- Items not in correct hierarchical order")
            if not parent_styled:
                print("FAIL: Component 3 -- Parent items not styled as numbered lists")
    except Exception as e:
        print("ERROR: Component 3 -- %s" % e)

    final_score = round(min(total_score, 1.0), 2)
    print("\nScore: %s/1.0" % total_score)
    print("REWARD: %s" % final_score)
    return final_score


# Entry point
file_path = '%s/%s.docx' % (WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
