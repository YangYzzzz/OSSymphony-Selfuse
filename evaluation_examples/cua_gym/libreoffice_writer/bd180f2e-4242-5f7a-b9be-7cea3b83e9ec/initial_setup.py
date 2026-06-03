"""
Initial Setup: Thesis document with front matter (pages 1-5) and Chapter 1 (page 6+)
Task ID: writer_acad_029
Domain: libreoffice_writer

Creates a multi-page thesis document with:
- Page 1: Title page
- Page 2: Abstract
- Page 3: Acknowledgments
- Page 4-5: Table of Contents
- Page 6+: Chapter 1
All pages use default continuous Arabic numbering in the footer.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_029'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section):
    """Add a simple Arabic page number in the footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Field code: PAGE
    r1 = fp.add_run()
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    r2 = fp.add_run()
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    r3 = fp.add_run()
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)

    # Style the footer text
    for run in fp.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'


def add_filler_text(doc, lines, font_name='Times New Roman', font_size=12):
    """Add multiple paragraphs of text to fill a page."""
    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5


def create_initial():
    doc = Document()

    # ----- Page Setup -----
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Add page number footer to first section (continuous Arabic numbering)
    add_page_number_footer(section)

    # ===== PAGE 1: TITLE PAGE =====
    # Add vertical space before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Machine Learning Approaches for\nPredictive Analysis of Climate Change\nPatterns in Southeast Asia')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author.add_run('by\n\nElena Vasquez Rodriguez')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacer

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = dept.add_run('Department of Environmental Data Science\nFaculty of Science and Technology\nUniversity of Singapore\n\nMarch 2025')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ===== PAGE BREAK -> PAGE 2: ABSTRACT =====
    doc.add_page_break()

    heading = doc.add_heading('Abstract', level=1)
    for run in heading.runs:
        run.font.name = 'Times New Roman'

    abstract_text = [
        'This thesis investigates the application of advanced machine learning techniques '
        'to predict and analyze climate change patterns across Southeast Asia, a region '
        'particularly vulnerable to environmental disruption. Drawing on satellite imagery, '
        'meteorological station data, and ocean buoy measurements collected between 2005 and 2024, '
        'we develop a suite of deep learning models capable of forecasting temperature anomalies, '
        'precipitation shifts, and extreme weather event frequencies at sub-regional granularity.',

        'Our primary contribution is the development of ClimateNet-SEA, a transformer-based '
        'architecture that integrates spatiotemporal features from heterogeneous climate data '
        'sources. The model achieves a mean absolute error of 0.34 degrees Celsius for '
        'twelve-month temperature forecasts, representing a 23% improvement over existing '
        'state-of-the-art methods. Additionally, we introduce a novel attention mechanism '
        'that captures long-range dependencies in monsoon circulation patterns.',

        'We validate our approach across eleven countries in the ASEAN region, demonstrating '
        'consistent performance improvements in both tropical maritime and continental climate '
        'zones. The model successfully identifies emerging climate tipping points in the '
        'Mekong Delta and Coral Triangle regions, providing actionable insights for '
        'policymakers and environmental agencies.',

        'Keywords: machine learning, climate change prediction, deep learning, Southeast Asia, '
        'spatiotemporal analysis, transformer networks, environmental data science',
    ]
    add_filler_text(doc, abstract_text)

    # ===== PAGE BREAK -> PAGE 3: ACKNOWLEDGMENTS =====
    doc.add_page_break()

    heading = doc.add_heading('Acknowledgments', level=1)
    for run in heading.runs:
        run.font.name = 'Times New Roman'

    ack_text = [
        'I would like to express my deepest gratitude to my primary supervisor, '
        'Professor Hiroshi Tanaka, for his unwavering guidance, insightful critiques, and '
        'endless patience throughout this research journey. His expertise in computational '
        'climate science has been instrumental in shaping this work.',

        'I am equally thankful to my co-supervisor, Dr. Amara Okonkwo, whose expertise in '
        'machine learning architectures provided critical direction during the model '
        'development phase. Her mentorship extended well beyond the academic sphere, and '
        'I am profoundly grateful for her encouragement during challenging periods.',

        'Special thanks to the members of the Climate Intelligence Lab at the University of '
        'Singapore, particularly Dr. Wei Lin Chen and Dr. Priya Nair, for countless '
        'discussions that refined my understanding of monsoon dynamics and ocean-atmosphere '
        'coupling mechanisms.',

        'I gratefully acknowledge the financial support provided by the ASEAN Research '
        'Foundation through Grant No. ARF-2022-0847 and the Singapore National Research '
        'Foundation Climate Impact Programme. Without this funding, the extensive data '
        'collection and computational resources required for this work would not have '
        'been possible.',

        'The satellite data used in this study was generously provided by the European '
        'Space Agency Copernicus Programme and the Japan Aerospace Exploration Agency. '
        'I thank their data distribution teams for prompt and helpful responses to '
        'our data requests.',

        'Finally, I dedicate this thesis to my family: my parents, Carlos and Maria, '
        'whose sacrifices made my education possible; my partner, James, for his steadfast '
        'support and understanding during late nights and missed weekends; and my daughter, '
        'Sofia, who arrived midway through this PhD and gave me the strongest motivation '
        'to see it through to completion.',
    ]
    add_filler_text(doc, ack_text)

    # ===== PAGE BREAK -> PAGE 4-5: TABLE OF CONTENTS =====
    doc.add_page_break()

    heading = doc.add_heading('Table of Contents', level=1)
    for run in heading.runs:
        run.font.name = 'Times New Roman'

    toc_entries = [
        ('Abstract', 'ii'),
        ('Acknowledgments', 'iii'),
        ('Table of Contents', 'iv'),
        ('List of Figures', 'vi'),
        ('List of Tables', 'viii'),
        ('List of Abbreviations', 'ix'),
        ('', ''),
        ('Chapter 1: Introduction', '1'),
        ('   1.1  Background and Motivation', '1'),
        ('   1.2  Climate Change in Southeast Asia: An Overview', '4'),
        ('   1.3  The Role of Machine Learning in Climate Science', '8'),
        ('   1.4  Research Objectives', '12'),
        ('   1.5  Thesis Structure', '14'),
        ('', ''),
        ('Chapter 2: Literature Review', '16'),
        ('   2.1  Traditional Climate Modeling Approaches', '16'),
        ('   2.2  Machine Learning for Environmental Prediction', '22'),
        ('   2.3  Deep Learning in Spatiotemporal Analysis', '28'),
        ('   2.4  Transfer Learning for Climate Applications', '35'),
        ('   2.5  Gaps in Current Research', '40'),
        ('', ''),
        ('Chapter 3: Methodology', '43'),
        ('   3.1  Data Collection and Preprocessing', '43'),
        ('   3.2  ClimateNet-SEA Architecture', '51'),
        ('   3.3  Attention Mechanism Design', '58'),
        ('   3.4  Training Procedure and Hyperparameters', '63'),
        ('   3.5  Evaluation Metrics', '67'),
        ('', ''),
        ('Chapter 4: Results and Analysis', '70'),
        ('   4.1  Temperature Anomaly Predictions', '70'),
        ('   4.2  Precipitation Pattern Forecasting', '78'),
        ('   4.3  Extreme Weather Event Detection', '85'),
        ('   4.4  Regional Performance Comparison', '91'),
        ('   4.5  Ablation Studies', '97'),
        ('', ''),
        ('Chapter 5: Discussion', '102'),
        ('   5.1  Interpretation of Key Findings', '102'),
        ('   5.2  Policy Implications', '108'),
        ('   5.3  Limitations and Future Work', '113'),
        ('', ''),
        ('Chapter 6: Conclusion', '118'),
        ('', ''),
        ('References', '122'),
        ('', ''),
        ('Appendix A: Supplementary Figures', '138'),
        ('Appendix B: Model Configuration Details', '145'),
        ('Appendix C: Data Source Catalogue', '150'),
    ]

    for entry, page in toc_entries:
        if not entry:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(1)

        # Add tab stop for right-aligned page number with dot leader
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(5.75), alignment=2, leader=2)  # RIGHT, DOTS

        run = para.add_run(f'{entry}\t{page}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # ===== PAGE BREAK -> PAGE 6: CHAPTER 1 =====
    doc.add_page_break()

    heading = doc.add_heading('Chapter 1: Introduction', level=1)
    for run in heading.runs:
        run.font.name = 'Times New Roman'

    ch1_text = [
        '1.1  Background and Motivation',
        '',
        'The accelerating pace of climate change represents one of the most pressing '
        'challenges facing humanity in the twenty-first century. Across the globe, rising '
        'temperatures, shifting precipitation patterns, and increasing frequencies of extreme '
        'weather events are reshaping ecosystems, threatening food security, and displacing '
        'communities. Among the regions most acutely affected by these changes is Southeast '
        'Asia, a geographically diverse area encompassing eleven nations and over 680 million '
        'people.',

        'Southeast Asia occupies a unique position in the global climate system. Situated '
        'between the Pacific and Indian Oceans, the region is influenced by complex '
        'ocean-atmosphere interactions, including the El Nino-Southern Oscillation (ENSO), '
        'the Indian Ocean Dipole (IOD), and the Asian monsoon system. These interconnected '
        'climate drivers produce highly variable weather patterns that challenge traditional '
        'forecasting methodologies.',

        'The economic implications of climate change in Southeast Asia are staggering. '
        'The Asian Development Bank estimates that without significant adaptation measures, '
        'the region could experience GDP losses of up to 11 percent by 2100. Agricultural '
        'productivity, which employs over 30 percent of the workforce in countries such as '
        'Myanmar, Cambodia, and Vietnam, is particularly vulnerable to temperature increases '
        'and altered monsoon timing.',

        'In this context, the ability to accurately predict climate trends at regional and '
        'sub-regional scales becomes not merely an academic pursuit but a practical necessity '
        'for effective policy planning and resource allocation. Traditional numerical weather '
        'prediction models, while valuable, often lack the spatial resolution and computational '
        'efficiency required for localized climate projections in data-sparse regions.',

        'Recent advances in machine learning, particularly in deep learning and transformer '
        'architectures, offer promising new avenues for climate prediction. These data-driven '
        'approaches can learn complex nonlinear relationships from large datasets, potentially '
        'complementing or even surpassing physics-based models in certain forecasting tasks. '
        'However, the application of these techniques to Southeast Asian climate data remains '
        'relatively unexplored, presenting a significant research gap that this thesis aims '
        'to address.',

        '1.2  Climate Change in Southeast Asia: An Overview',
        '',
        'Southeast Asia has experienced a mean temperature increase of approximately 0.14 to '
        '0.20 degrees Celsius per decade since 1960, with some sub-regions, particularly the '
        'Mekong River Basin and the Indonesian archipelago, exhibiting even higher warming '
        'rates. This warming trend is accompanied by measurable changes in monsoon onset '
        'timing, with delays of 5 to 15 days observed across multiple monitoring stations.',

        'Precipitation patterns have become increasingly erratic. While total annual '
        'rainfall in some areas has remained relatively stable, the distribution has shifted '
        'markedly, with longer dry periods punctuated by more intense rainfall events. In '
        'the Philippines alone, the frequency of Category 4 and 5 typhoons has increased by '
        'approximately 18 percent over the past four decades.',

        'Sea level rise poses an existential threat to low-lying coastal areas and island '
        'nations within the region. The Intergovernmental Panel on Climate Change projects '
        'a rise of 0.4 to 0.8 meters by 2100 under moderate emissions scenarios, which '
        'would directly affect the approximately 77 million people living in coastal areas '
        'less than one meter above current sea level.',
    ]

    for text in ch1_text:
        if text == '':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue
        if text.startswith('1.'):
            heading_para = doc.add_heading(text, level=2)
            for run in heading_para.runs:
                run.font.name = 'Times New Roman'
        else:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = Inches(0.5)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
