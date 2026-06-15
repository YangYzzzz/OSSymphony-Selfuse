"""
Reward Script: Set up different page numbering for thesis
Task ID: writer_acad_029
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Document has at least 2 sections (section break exists)
  Component 2 (0.30): First section uses lowerRoman page numbering format
  Component 3 (0.25): Second section uses decimal page numbering with start=1
  Component 4 (0.20): Second section footer has PAGE field and is not linked to previous
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_029'


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(1.0)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires:
    - Pages 1-5 (front matter): lowercase Roman numeral page numbering (i, ii, iii...)
    - Page 6+ (Chapter 1): Arabic numerals starting from 1

    This requires a section break between front matter and Chapter 1,
    with distinct pgNumType settings on each section.
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    num_sections = len(doc.sections)

    # Component 1: Document has at least 2 sections (0.25 points)
    # Initial doc has 1 section; golden has 2 (section break inserted before Chapter 1)
    try:
        if num_sections >= 2:
            print(f"PASS: Component 1 -- Document has {num_sections} sections (>=2) (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 -- Document has {num_sections} section(s), expected >=2")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: First section uses lowerRoman page numbering (0.30 points)
    # Initial doc has no pgNumType; golden has fmt=lowerRoman
    try:
        if num_sections >= 1:
            sect0 = doc.sections[0]._sectPr
            pgNumType0 = sect0.find(qn('w:pgNumType'))
            if pgNumType0 is not None:
                fmt0 = pgNumType0.get(qn('w:fmt'))
                if fmt0 == 'lowerRoman':
                    print(f"PASS: Component 2 -- Section 0 pgNumType fmt=lowerRoman (0.30 pts)")
                    total_score += 0.30
                else:
                    print(f"FAIL: Component 2 -- Section 0 pgNumType fmt={fmt0}, expected lowerRoman")
            else:
                print("FAIL: Component 2 -- Section 0 has no pgNumType element")
        else:
            print("FAIL: Component 2 -- No sections to check")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Second section uses decimal numbering with start=1 (0.25 points)
    # Initial doc has no second section; golden has fmt=decimal, start=1
    try:
        if num_sections >= 2:
            sect1 = doc.sections[1]._sectPr
            pgNumType1 = sect1.find(qn('w:pgNumType'))
            if pgNumType1 is not None:
                fmt1 = pgNumType1.get(qn('w:fmt'))
                start1 = pgNumType1.get(qn('w:start'))
                fmt_ok = fmt1 == 'decimal' or fmt1 is None  # decimal is also default
                start_ok = start1 == '1'
                if fmt_ok and start_ok:
                    print(f"PASS: Component 3 -- Section 1 pgNumType fmt={fmt1}, start={start1} (0.25 pts)")
                    total_score += 0.25
                else:
                    print(f"FAIL: Component 3 -- Section 1 pgNumType fmt={fmt1}, start={start1}; expected decimal/start=1")
            else:
                print("FAIL: Component 3 -- Section 1 has no pgNumType element")
        else:
            print("FAIL: Component 3 -- Fewer than 2 sections, cannot check section 1")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Second section has its own footer with PAGE field, not linked to previous (0.20 points)
    # This ensures page numbers are actually displayed in section 2 independently
    try:
        if num_sections >= 2:
            sect1 = doc.sections[1]
            footer = sect1.footer
            is_linked = footer.is_linked_to_previous
            # Check for PAGE field code in footer XML
            page_field_count = 0
            if footer and footer.paragraphs:
                for fp in footer.paragraphs:
                    for run_elem in fp._element.findall(qn('w:r')):
                        for child in run_elem:
                            if child.tag == qn('w:instrText'):
                                if 'PAGE' in (child.text or ''):
                                    page_field_count += 1
            has_page_field = page_field_count > 0
            if not is_linked and has_page_field:
                print(f"PASS: Component 4 -- Section 1 footer: not linked, has PAGE field (0.20 pts)")
                total_score += 0.20
            elif is_linked and has_page_field:
                # Linked to previous but inherits PAGE field -- partial: the numbering
                # format change in pgNumType will still take effect even if linked
                print(f"PARTIAL: Component 4 -- Section 1 footer linked to previous but has PAGE field (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 -- Section 1 footer: linked={is_linked}, has_page_field={has_page_field}")
        else:
            print("FAIL: Component 4 -- Fewer than 2 sections")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
