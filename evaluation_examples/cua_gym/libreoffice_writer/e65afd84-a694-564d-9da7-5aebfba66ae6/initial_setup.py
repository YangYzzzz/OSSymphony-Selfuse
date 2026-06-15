"""
Initial Setup: Software Requirements Specification with manually typed requirement IDs
Task ID: writer_tech_048
Domain: libreoffice_writer

Creates a realistic software requirements specification document with 8 requirement
paragraphs that have manually typed REQ-001 through REQ-008 prefixes (plain text,
no automatic list numbering).
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_048'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Title ---
    title = doc.add_heading('Software Requirements Specification', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Subtitle ---
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('CloudSync Platform v2.4')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # --- Document info ---
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = info.add_run('Prepared by: Systems Engineering Division\nRelease Date: March 15, 2025\nDocument Version: 1.3')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # spacing

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This document outlines the functional and non-functional requirements for the '
        'CloudSync Platform version 2.4. The platform provides enterprise-grade file '
        'synchronization, real-time collaboration, and access control capabilities across '
        'distributed teams. All requirements listed below have been reviewed and approved '
        'by the Product Review Board on February 28, 2025.'
    )

    # --- Section 2: Scope ---
    doc.add_heading('2. Scope', level=1)
    doc.add_paragraph(
        'The CloudSync Platform serves organizations with 50 to 10,000 users, supporting '
        'Windows, macOS, Linux, iOS, and Android clients. The backend infrastructure runs '
        'on AWS with multi-region failover. This specification covers the core synchronization '
        'engine, user authentication module, administrative dashboard, and API gateway.'
    )

    # --- Section 3: Functional Requirements ---
    doc.add_heading('3. Functional Requirements', level=1)
    doc.add_paragraph(
        'The following requirements define the expected behavior of the CloudSync Platform. '
        'Each requirement is assigned a unique identifier for traceability throughout the '
        'development lifecycle.'
    )

    # 8 requirement paragraphs with manually typed IDs (plain text, NO list style)
    requirements = [
        ('REQ-001:', 'The system shall support simultaneous file synchronization across '
         'a minimum of 500 concurrent users without degradation in transfer speeds below '
         '10 MB/s per session on a standard enterprise network connection.'),

        ('REQ-002:', 'User authentication shall integrate with SAML 2.0, OAuth 2.0, and '
         'LDAP directory services, allowing single sign-on (SSO) across all supported '
         'client platforms with session tokens expiring after 8 hours of inactivity.'),

        ('REQ-003:', 'The conflict resolution engine shall detect and resolve file version '
         'conflicts within 3 seconds of detection, presenting the user with a merge dialog '
         'that displays both versions side-by-side with highlighted differences.'),

        ('REQ-004:', 'Administrative users shall be able to configure role-based access '
         'control (RBAC) policies with a minimum of 12 predefined roles and the ability '
         'to create up to 50 custom roles per organizational unit.'),

        ('REQ-005:', 'The platform shall maintain a complete audit trail of all file '
         'operations including uploads, downloads, deletions, sharing events, and permission '
         'changes, retaining logs for a configurable period of 90 to 365 days.'),

        ('REQ-006:', 'Real-time collaboration features shall support co-editing of documents '
         'by up to 25 simultaneous users with cursor position tracking, change attribution, '
         'and automatic save intervals of no more than 5 seconds.'),

        ('REQ-007:', 'The API gateway shall expose RESTful endpoints with OpenAPI 3.0 '
         'documentation, supporting rate limiting of 1,000 requests per minute per API key '
         'and returning responses within 200 milliseconds for 95% of read operations.'),

        ('REQ-008:', 'Data encryption shall use AES-256 for files at rest and TLS 1.3 for '
         'data in transit, with automatic key rotation every 90 days managed through an '
         'integrated key management service compatible with AWS KMS and Azure Key Vault.'),
    ]

    for req_id, req_text in requirements:
        para = doc.add_paragraph()
        # Manually typed ID as bold run
        id_run = para.add_run(f'{req_id} ')
        id_run.bold = True
        id_run.font.size = Pt(11)
        # Requirement text as normal run
        text_run = para.add_run(req_text)
        text_run.font.size = Pt(11)

    # --- Section 4: Non-Functional Requirements ---
    doc.add_heading('4. Non-Functional Requirements', level=1)
    doc.add_paragraph(
        'System availability shall maintain 99.95% uptime measured monthly, with planned '
        'maintenance windows limited to 4 hours per month during off-peak hours (Saturday '
        '02:00-06:00 UTC). Disaster recovery procedures shall achieve a Recovery Time '
        'Objective (RTO) of 15 minutes and Recovery Point Objective (RPO) of 5 minutes.'
    )
    doc.add_paragraph(
        'Performance benchmarks require the platform to handle peak loads of 10,000 '
        'simultaneous file operations with average latency under 150 milliseconds. The '
        'system shall scale horizontally to accommodate 200% traffic spikes within 60 '
        'seconds through auto-scaling policies configured in the deployment infrastructure.'
    )

    # --- Section 5: Approval ---
    doc.add_heading('5. Approval', level=1)
    doc.add_paragraph(
        'This specification has been reviewed and approved by the following stakeholders:'
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['Name', 'Role', 'Date']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    approvers = [
        ['Dr. Elena Vasquez', 'VP of Engineering', '2025-02-28'],
        ['James Thornton', 'Product Director', '2025-02-28'],
        ['Priya Krishnamurthy', 'QA Lead', '2025-03-01'],
    ]
    for r, row_data in enumerate(approvers, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
