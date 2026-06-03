"""
Initial Setup: Machine Learning Handbook - header same on all pages
Task ID: writer_page_078
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKDIR = '/home/user'
TASK_ID = 'writer_page_078'
OUTPUT = f'{WORKDIR}/Desktop/ml_handbook.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading(doc, text, level=1):
    """Add a heading paragraph."""
    para = doc.add_heading(text, level=level)
    return para


def add_paragraph(doc, text):
    """Add a normal paragraph."""
    para = doc.add_paragraph(text)
    return para


def create_initial():
    doc = Document()

    # --- Section setup: A4 portrait, mirrored margins ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Mirrored margins: set gutter/mirror via XML
    # inner margin = 2.5cm, outer = 2.0cm, top = 2.5cm, bottom = 2.5cm
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    # For mirrored margins: left becomes inner (2.5cm), right becomes outer (2.0cm)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # Enable mirrored margins via XML
    pgMar = section._sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        section._sectPr.append(pgMar)
    pgMar.set(qn('w:mirrorMargins'), '1')

    # Enable header/footer distance
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    # --- Header: same on all pages (NOT different odd/even) ---
    # Ensure different_odd_even_pages is NOT set (same on all pages)
    # Remove evenAndOddHeaders element if present
    titlePg = section._sectPr.find(qn('w:titlePg'))
    # Do NOT set different odd/even pages header
    evenOdd = section._sectPr.find(qn('w:evenAndOddHeaders'))
    if evenOdd is not None:
        section._sectPr.remove(evenOdd)

    header = section.header
    header.is_linked_to_previous = False

    # Clear existing paragraphs and set header text
    for para in header.paragraphs:
        para.clear()

    header_para = header.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header_para.add_run('Machine Learning Handbook')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # --- Document content: 20-page machine learning handbook chapter ---
    # Title
    title = doc.add_heading('Machine Learning Handbook', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('Part II: Advanced Topics', level=1)

    doc.add_heading('Chapter 5: Neural Networks and Deep Learning', level=2)

    doc.add_paragraph(
        'Neural networks are computational models inspired by the human brain\'s structure. '
        'They consist of interconnected nodes (neurons) organized in layers: an input layer, '
        'one or more hidden layers, and an output layer. Each connection has an associated '
        'weight that is adjusted during training to minimize prediction error.'
    )

    doc.add_paragraph(
        'Deep learning refers to neural networks with many hidden layers. These deep architectures '
        'can learn increasingly abstract representations of data. The breakthrough in deep learning '
        'came around 2012 when convolutional neural networks achieved unprecedented performance '
        'on image classification benchmarks such as ImageNet.'
    )

    doc.add_heading('5.1 Activation Functions', level=3)

    doc.add_paragraph(
        'Activation functions introduce non-linearity into neural networks, enabling them to '
        'learn complex patterns. Common activation functions include:'
    )

    doc.add_paragraph('Sigmoid: f(x) = 1 / (1 + e^(-x)). Range: (0, 1). Used in binary classification output layers.', style='List Bullet')
    doc.add_paragraph('ReLU (Rectified Linear Unit): f(x) = max(0, x). Most widely used in hidden layers due to computational efficiency.', style='List Bullet')
    doc.add_paragraph('Tanh: f(x) = (e^x - e^(-x)) / (e^x + e^(-x)). Range: (-1, 1). Zero-centered, often preferred over sigmoid.', style='List Bullet')
    doc.add_paragraph('Softmax: Converts raw scores to probabilities. Used in multi-class classification output layers.', style='List Bullet')

    doc.add_paragraph(
        'The choice of activation function significantly affects the training dynamics and the '
        'network\'s ability to converge. Vanishing gradient problems with sigmoid and tanh '
        'led to widespread adoption of ReLU and its variants (Leaky ReLU, ELU, SELU).'
    )

    doc.add_page_break()

    doc.add_heading('5.2 Backpropagation', level=3)

    doc.add_paragraph(
        'Backpropagation is the algorithm used to train neural networks. It computes the gradient '
        'of the loss function with respect to each weight by applying the chain rule of calculus. '
        'The algorithm proceeds in two phases:'
    )

    doc.add_paragraph('Forward pass: Input data flows through the network, producing predictions.', style='List Number')
    doc.add_paragraph('Backward pass: The error is propagated backwards through the network, computing gradients.', style='List Number')
    doc.add_paragraph('Weight update: Weights are adjusted in the direction that reduces the loss.', style='List Number')

    doc.add_paragraph(
        'Stochastic Gradient Descent (SGD) and its variants (Adam, RMSProp, AdaGrad) are the '
        'most common optimization algorithms used with backpropagation. The learning rate is a '
        'critical hyperparameter that controls the step size during weight updates.'
    )

    doc.add_heading('5.3 Convolutional Neural Networks (CNNs)', level=3)

    doc.add_paragraph(
        'Convolutional Neural Networks are specialized architectures designed for processing '
        'grid-structured data such as images. They leverage spatial locality and translational '
        'invariance through shared weights in convolutional filters.'
    )

    doc.add_paragraph(
        'Key components of CNNs include convolutional layers (feature extraction), pooling layers '
        '(spatial downsampling), and fully connected layers (classification). Architectures like '
        'VGG, ResNet, Inception, and EfficientNet have pushed the state-of-the-art on computer '
        'vision benchmarks.'
    )

    doc.add_page_break()

    doc.add_heading('Chapter 6: Recurrent Neural Networks and Transformers', level=2)

    doc.add_paragraph(
        'Recurrent Neural Networks (RNNs) are designed to process sequential data by maintaining '
        'a hidden state that captures information about previous inputs. This makes them suitable '
        'for natural language processing, time series analysis, and speech recognition.'
    )

    doc.add_heading('6.1 Long Short-Term Memory (LSTM)', level=3)

    doc.add_paragraph(
        'LSTMs were introduced by Hochreiter and Schmidhuber (1997) to address the vanishing '
        'gradient problem in standard RNNs. They use gating mechanisms to selectively remember '
        'or forget information over long sequences:'
    )

    doc.add_paragraph('Forget gate: Decides what information to discard from the cell state.', style='List Bullet')
    doc.add_paragraph('Input gate: Decides which new information to store in the cell state.', style='List Bullet')
    doc.add_paragraph('Output gate: Decides what to output based on the cell state.', style='List Bullet')

    doc.add_paragraph(
        'LSTMs have been successfully applied to machine translation, sentiment analysis, '
        'music generation, and many other sequence modeling tasks. They remain a strong '
        'baseline despite the rise of transformer architectures.'
    )

    doc.add_page_break()

    doc.add_heading('6.2 The Transformer Architecture', level=3)

    doc.add_paragraph(
        'The Transformer architecture, introduced in "Attention Is All You Need" (Vaswani et al., '
        '2017), revolutionized natural language processing. Unlike RNNs, Transformers process '
        'all tokens in parallel using self-attention mechanisms, making them highly efficient '
        'on modern hardware.'
    )

    doc.add_paragraph(
        'The core innovation is the multi-head self-attention mechanism, which allows the model '
        'to attend to different parts of the input sequence simultaneously. This enables the '
        'capture of long-range dependencies without the sequential bottleneck of RNNs.'
    )

    doc.add_heading('6.3 BERT and GPT Models', level=3)

    doc.add_paragraph(
        'BERT (Bidirectional Encoder Representations from Transformers) is a pre-trained language '
        'model that uses masked language modeling to learn bidirectional context representations. '
        'It achieves state-of-the-art results on many NLP tasks through fine-tuning.'
    )

    doc.add_paragraph(
        'GPT (Generative Pre-trained Transformer) models use autoregressive language modeling, '
        'predicting the next token given previous tokens. GPT-3 and GPT-4 demonstrated emergent '
        'capabilities in few-shot learning across diverse tasks without task-specific fine-tuning.'
    )

    doc.add_page_break()

    doc.add_heading('Chapter 7: Reinforcement Learning', level=2)

    doc.add_paragraph(
        'Reinforcement Learning (RL) is a paradigm where an agent learns to make decisions by '
        'interacting with an environment to maximize cumulative reward. Unlike supervised learning, '
        'RL does not require labeled examples; instead, it learns from delayed feedback signals.'
    )

    doc.add_heading('7.1 Markov Decision Processes', level=3)

    doc.add_paragraph(
        'The formal framework for RL is the Markov Decision Process (MDP), defined by:'
    )

    doc.add_paragraph('State space S: All possible states of the environment.', style='List Bullet')
    doc.add_paragraph('Action space A: All possible actions the agent can take.', style='List Bullet')
    doc.add_paragraph('Transition function P(s\'|s,a): Probability of transitioning to state s\' from state s after action a.', style='List Bullet')
    doc.add_paragraph('Reward function R(s,a): Immediate reward for taking action a in state s.', style='List Bullet')
    doc.add_paragraph('Discount factor γ: Balances immediate vs. future rewards (0 ≤ γ ≤ 1).', style='List Bullet')

    doc.add_page_break()

    doc.add_heading('7.2 Q-Learning and Deep Q-Networks', level=3)

    doc.add_paragraph(
        'Q-Learning is a model-free RL algorithm that learns the optimal action-value function '
        'Q(s, a), representing the expected cumulative reward of taking action a in state s and '
        'following the optimal policy thereafter.'
    )

    doc.add_paragraph(
        'Deep Q-Networks (DQN) combine Q-Learning with deep neural networks to handle '
        'high-dimensional state spaces. Key innovations include experience replay (storing '
        'and sampling past transitions) and target networks (stabilizing training by '
        'maintaining a periodically updated copy of the Q-network).'
    )

    doc.add_heading('7.3 Policy Gradient Methods', level=3)

    doc.add_paragraph(
        'Policy gradient methods directly optimize the policy π(a|s) rather than learning '
        'a value function. The REINFORCE algorithm uses Monte Carlo sampling to estimate '
        'policy gradients. Actor-critic methods combine value function estimation with '
        'policy optimization for reduced variance.'
    )

    doc.add_paragraph(
        'Proximal Policy Optimization (PPO) and Trust Region Policy Optimization (TRPO) '
        'are state-of-the-art policy gradient algorithms that constrain policy updates '
        'to maintain training stability. PPO is widely used due to its simplicity and '
        'strong empirical performance across diverse tasks.'
    )

    doc.add_page_break()

    doc.add_heading('Chapter 8: Unsupervised and Self-Supervised Learning', level=2)

    doc.add_paragraph(
        'Unsupervised learning algorithms discover patterns in data without labeled examples. '
        'Key approaches include clustering (k-means, DBSCAN, hierarchical clustering), '
        'dimensionality reduction (PCA, t-SNE, UMAP), and generative modeling.'
    )

    doc.add_heading('8.1 Generative Adversarial Networks', level=3)

    doc.add_paragraph(
        'Generative Adversarial Networks (GANs), introduced by Goodfellow et al. (2014), '
        'consist of two neural networks in adversarial competition: a generator that creates '
        'synthetic data and a discriminator that distinguishes real from synthetic data.'
    )

    doc.add_paragraph(
        'The training process resembles a minimax game: the generator improves to fool the '
        'discriminator, while the discriminator improves to detect fakes. This adversarial '
        'training has produced remarkable results in image synthesis, style transfer, '
        'and data augmentation.'
    )

    doc.add_page_break()

    doc.add_heading('8.2 Variational Autoencoders', level=3)

    doc.add_paragraph(
        'Variational Autoencoders (VAEs) are probabilistic generative models that learn a '
        'latent representation of data. They consist of an encoder network that maps inputs '
        'to a distribution in latent space, and a decoder network that reconstructs inputs '
        'from sampled latent vectors.'
    )

    doc.add_paragraph(
        'VAEs are trained with a combination of reconstruction loss and KL divergence regularization, '
        'ensuring the latent space is continuous and suitable for generation. Unlike GANs, '
        'VAEs provide explicit probabilistic inference and are more stable to train.'
    )

    doc.add_heading('8.3 Self-Supervised Learning', level=3)

    doc.add_paragraph(
        'Self-supervised learning leverages the structure of unlabeled data to create supervised '
        'training signals. In computer vision, methods like SimCLR and MoCo learn representations '
        'by maximizing agreement between differently augmented views of the same image.'
    )

    doc.add_paragraph(
        'Masked language modeling (used in BERT) is a form of self-supervised learning for text. '
        'The model predicts masked tokens from context, learning rich language representations '
        'without any human-labeled data. These pre-trained representations transfer effectively '
        'to downstream tasks with minimal fine-tuning.'
    )

    doc.add_page_break()

    doc.add_heading('Chapter 9: Model Evaluation and Validation', level=2)

    doc.add_paragraph(
        'Rigorous evaluation is essential for assessing the true performance of machine learning '
        'models. Overfitting, data leakage, and distribution shift are common pitfalls that '
        'can lead to misleadingly optimistic performance estimates.'
    )

    doc.add_heading('9.1 Cross-Validation', level=3)

    doc.add_paragraph(
        'k-Fold cross-validation partitions the dataset into k equal folds, training on k-1 folds '
        'and evaluating on the remaining fold. This is repeated k times, and results are averaged. '
        'Stratified k-fold ensures each fold maintains the class distribution of the full dataset.'
    )

    doc.add_heading('9.2 Performance Metrics', level=3)

    doc.add_paragraph(
        'Classification metrics include accuracy, precision, recall, F1-score, and AUC-ROC. '
        'Regression metrics include mean squared error (MSE), mean absolute error (MAE), and '
        'R-squared. Choosing the appropriate metric depends on the problem requirements and '
        'the relative cost of different types of errors.'
    )

    doc.add_page_break()

    doc.add_heading('Chapter 10: Practical Considerations', level=2)

    doc.add_paragraph(
        'Deploying machine learning models in production requires careful consideration of '
        'computational efficiency, latency requirements, fairness, and robustness to '
        'distribution shift. Model compression techniques such as pruning, quantization, '
        'and knowledge distillation can significantly reduce inference costs.'
    )

    doc.add_heading('10.1 Hyperparameter Tuning', level=3)

    doc.add_paragraph(
        'Hyperparameter tuning is the process of finding the optimal configuration for a '
        'machine learning model. Methods include grid search, random search, Bayesian '
        'optimization, and neural architecture search (NAS). AutoML frameworks automate '
        'much of this process.'
    )

    doc.add_heading('10.2 Transfer Learning', level=3)

    doc.add_paragraph(
        'Transfer learning leverages knowledge from pre-trained models to improve performance '
        'on new tasks with limited data. Fine-tuning pre-trained models (e.g., ImageNet-trained '
        'CNNs, BERT) on domain-specific data is one of the most effective techniques in '
        'applied machine learning.'
    )

    doc.add_paragraph(
        'Domain adaptation addresses the challenge of distributional shift between source and '
        'target domains. Techniques include feature alignment, adversarial training, and '
        'gradual fine-tuning to prevent catastrophic forgetting of source domain knowledge.'
    )

    doc.add_page_break()

    doc.add_heading('Appendix A: Mathematical Foundations', level=1)

    doc.add_paragraph(
        'This appendix reviews key mathematical concepts underlying modern machine learning: '
        'linear algebra (matrices, eigendecomposition, SVD), probability theory (Bayes rule, '
        'distributions, MLE), calculus (gradients, chain rule, optimization), and information '
        'theory (entropy, mutual information, KL divergence).'
    )

    doc.add_heading('A.1 Linear Algebra Review', level=3)

    doc.add_paragraph(
        'Vectors and matrices are the fundamental data structures in machine learning. A vector '
        'x ∈ ℝ^n represents a point in n-dimensional space. A matrix A ∈ ℝ^(m×n) represents '
        'a linear transformation from ℝ^n to ℝ^m. Matrix multiplication, transpose, inverse, '
        'and decompositions (LU, QR, SVD, eigendecomposition) are essential operations.'
    )

    doc.add_heading('A.2 Probability and Statistics', level=3)

    doc.add_paragraph(
        'Machine learning is fundamentally probabilistic. Key concepts include probability '
        'distributions (Gaussian, Bernoulli, Categorical, Dirichlet), expectation and variance, '
        'conditional probability, Bayes\' theorem, maximum likelihood estimation (MLE), '
        'and maximum a posteriori (MAP) estimation.'
    )

    doc.add_page_break()

    doc.add_heading('Appendix B: Common Datasets and Benchmarks', level=1)

    doc.add_paragraph(
        'The machine learning community has established standard benchmarks for comparing '
        'algorithm performance. This appendix describes key datasets used throughout '
        'the handbook examples.'
    )

    # Add a table of datasets
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dataset'
    hdr_cells[1].text = 'Task'
    hdr_cells[2].text = 'Size'
    hdr_cells[3].text = 'Modality'

    datasets = [
        ('MNIST', 'Digit Classification', '70,000 images', 'Image'),
        ('CIFAR-10', 'Object Recognition', '60,000 images', 'Image'),
        ('ImageNet', 'Object Classification', '1.2M images', 'Image'),
        ('IMDb Reviews', 'Sentiment Analysis', '50,000 reviews', 'Text'),
        ('SQuAD 2.0', 'Question Answering', '150,000 Q&A pairs', 'Text'),
        ('Penn Treebank', 'Language Modeling', '1M tokens', 'Text'),
        ('OpenAI Gym', 'Reinforcement Learning', 'Various envs', 'Simulation'),
        ('UCI ML Repository', 'Various', '500+ datasets', 'Tabular'),
    ]

    for name, task, size, modality in datasets:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = task
        row_cells[2].text = size
        row_cells[3].text = modality

    # Ensure desktop directory exists
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the document in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
