"""
Reward Script: Set header to have different content on left and right pages
Task ID: writer_page_078
Domain: libreoffice_writer
Scoring:
  Component 1: evenAndOddHeaders is enabled in section properties (0.35 pts)
  Component 2: Left (even) page header = 'Part II: Advanced Topics' centered (0.35 pts)
  Component 3: Right (odd) page header = 'Machine Learning Handbook' centered
               AND evenAndOddHeaders is enabled (0.30 pts)
Total: 1.0
"""

import os
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_page_078'
FILE_PATH = f'{WORKDIR}/Desktop/ml_handbook.docx'


def persist_app_state():
    """Attempt to save any unsaved GUI edits before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.

    Task: Set the header to have different content on left and right pages.
    - Left pages (even) show 'Part II: Advanced Topics' centered
    - Right pages (odd) show 'Machine Learning Handbook' centered
    - 'Same content on left and right pages' option must be unchecked
      (i.e., evenAndOddHeaders must be enabled in the section XML)

    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        section = doc.sections[0]
    except Exception as e:
        print(f"CRITICAL: Cannot access section: {e}")
        print("REWARD: 0.0")
        return 0.0

    # ------------------------------------------------------------------
    # Component 1: evenAndOddHeaders is enabled (0.35 points)
    # This flag in section XML enables different headers for left/right pages.
    # In OOXML: <w:evenAndOddHeaders/> present in <w:sectPr>.
    # Initial state: NOT present. Golden state: present.
    # ------------------------------------------------------------------
    even_and_odd_enabled = False
    try:
        sectPr = section._sectPr
        evenAndOddHeaders = sectPr.find(qn('w:evenAndOddHeaders'))
        if evenAndOddHeaders is not None:
            even_and_odd_enabled = True
            print("PASS: Component 1 — evenAndOddHeaders is enabled in section XML (0.35 pts)")
            total_score += 0.35
        else:
            print("FAIL: Component 1 — evenAndOddHeaders not found in section XML (expected enabled for different left/right headers)")
    except Exception as e:
        print(f"ERROR: Component 1 — could not check evenAndOddHeaders: {e}")

    # ------------------------------------------------------------------
    # Component 2: Left (even) page header = 'Part II: Advanced Topics' centered (0.35 points)
    # When evenAndOddHeaders is enabled, even pages use the even_page_header.
    # In a mirrored-margins document, left pages are even pages.
    # Initial state: even_page_header is empty. Golden state: 'Part II: Advanced Topics' centered.
    # ------------------------------------------------------------------
    try:
        even_header = section.even_page_header
        if even_header and even_header.paragraphs:
            even_para = even_header.paragraphs[0]
            even_text = even_para.text.strip()
            even_align = even_para.paragraph_format.alignment

            expected_even_text = "Part II: Advanced Topics"
            expected_align = WD_PARAGRAPH_ALIGNMENT.CENTER

            text_ok = even_text == expected_even_text
            align_ok = even_align == expected_align

            if text_ok and align_ok:
                print(f"PASS: Component 2 — Left (even) page header = '{even_text}' centered (0.35 pts)")
                total_score += 0.35
            elif text_ok and not align_ok:
                print(f"FAIL: Component 2 — Left (even) page header text correct ('{even_text}') but alignment is {even_align}, expected CENTER")
            elif not text_ok and align_ok:
                print(f"FAIL: Component 2 — Left (even) page header alignment correct but text is '{even_text}', expected '{expected_even_text}'")
            else:
                print(f"FAIL: Component 2 — Left (even) page header text='{even_text}' align={even_align}, expected text='{expected_even_text}' align=CENTER")
        else:
            print("FAIL: Component 2 — even_page_header has no paragraphs or is empty")
    except Exception as e:
        print(f"ERROR: Component 2 — could not check even_page_header: {e}")

    # ------------------------------------------------------------------
    # Component 3: Right (odd) page header = 'Machine Learning Handbook' centered
    #              AND evenAndOddHeaders is enabled (0.30 points)
    # This component requires evenAndOddHeaders to be set to ensure the odd header
    # is functioning as a "right-page-only" header (not a universal header).
    # Without evenAndOddHeaders, the odd header text matching alone doesn't count
    # because it's the same-for-all-pages header (precondition, not task result).
    # Initial state: fails because evenAndOddHeaders is not set.
    # Golden state: passes because evenAndOddHeaders is set AND odd header is correct.
    # ------------------------------------------------------------------
    try:
        if not even_and_odd_enabled:
            print("FAIL: Component 3 — evenAndOddHeaders not enabled, so right-page-only header not properly configured")
        else:
            odd_header = section.header
            if odd_header and odd_header.paragraphs:
                odd_para = odd_header.paragraphs[0]
                odd_text = odd_para.text.strip()
                odd_align = odd_para.paragraph_format.alignment

                expected_odd_text = "Machine Learning Handbook"
                expected_align = WD_PARAGRAPH_ALIGNMENT.CENTER

                text_ok = odd_text == expected_odd_text
                align_ok = odd_align == expected_align

                if text_ok and align_ok:
                    print(f"PASS: Component 3 — Right (odd) page header = '{odd_text}' centered, evenAndOddHeaders enabled (0.30 pts)")
                    total_score += 0.30
                elif text_ok and not align_ok:
                    print(f"FAIL: Component 3 — Right (odd) page header text correct ('{odd_text}') but alignment is {odd_align}, expected CENTER")
                elif not text_ok and align_ok:
                    print(f"FAIL: Component 3 — Right (odd) page header alignment correct but text is '{odd_text}', expected '{expected_odd_text}'")
                else:
                    print(f"FAIL: Component 3 — Right (odd) page header text='{odd_text}' align={odd_align}, expected text='{expected_odd_text}' align=CENTER")
            else:
                print("FAIL: Component 3 — odd (right) page header has no paragraphs or is empty")
    except Exception as e:
        print(f"ERROR: Component 3 — could not check odd (right) page header: {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved GUI edits before verifying
persist_app_state()

if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
