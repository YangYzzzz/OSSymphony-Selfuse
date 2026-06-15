"""
Reward Script: Complex class schedule table with merged cells in LibreOffice Writer
Task ID: writer_tbl_050
Domain: libreoffice_writer
Scoring:
  Component 1: Table with 6 rows x 6 columns exists           — 0.20 pts
  Component 2: Row 1 headers ('', Mon, Tue, Wed, Thu, Fri)    — 0.20 pts
  Component 3: A2-A3 vertically merged containing 'Morning'   — 0.20 pts
  Component 4: A4-A5 vertically merged containing 'Afternoon' — 0.20 pts
  Component 5: Row 6 horizontally merged (gridSpan=6) 'Lunch' — 0.20 pts
  Total: 1.0
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'writer_tbl_050'
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

FILE_PATH = f'{WORKDIR}/Desktop/class_schedule.docx'


def get_actual_tcs(table, row_idx):
    """Return the actual <w:tc> elements in a row (not python-docx merged duplicates)."""
    row_elem = table.rows[row_idx]._tr
    return row_elem.findall('{%s}tc' % W)


def get_tc_vmerge(tc):
    """
    Return the vMerge status of a tc element:
      'restart'  — start of a vertical merge group
      'continue' — continuation cell (empty <w:vMerge/> element)
      None       — no vMerge element (not part of a vertical merge)
    """
    tcPr = tc.find('{%s}tcPr' % W)
    if tcPr is None:
        return None
    vmerge = tcPr.find('{%s}vMerge' % W)
    if vmerge is None:
        return None
    val = vmerge.get('{%s}val' % W)
    if val == 'restart':
        return 'restart'
    # No val attribute or val="" means continue
    return 'continue'


def get_tc_gridspan(tc):
    """Return the gridSpan value of a tc element, or 1 if not set."""
    tcPr = tc.find('{%s}tcPr' % W)
    if tcPr is None:
        return 1
    gridspan = tcPr.find('{%s}gridSpan' % W)
    if gridspan is None:
        return 1
    val = gridspan.get('{%s}val' % W)
    return int(val) if val else 1


def get_tc_text(tc):
    """Get all text within a tc element."""
    return ''.join(t.text or '' for t in tc.iter('{%s}t' % W)).strip()


def verify_task(file_path):
    """
    Verify task completion: class schedule table with merged cells.
    Returns a float between 0.0 and 1.0.
    """
    total_score = 0.0

    # Load document — gate check
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Gate: table must exist
    if len(doc.tables) == 0:
        print("FAIL: No table found in document — task not completed")
        print(f"\nScore: 0.0/1.0")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table dimensions — 6 rows x 6 columns (0.20 pts)
    try:
        num_rows = len(table.rows)
        # Get column count from row 0 actual tc elements
        row0_tcs = get_actual_tcs(table, 0)
        num_cols_row0 = len(row0_tcs)

        if num_rows == 6 and num_cols_row0 == 6:
            print(f"PASS: Component 1 — Table is 6 rows x 6 cols (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 1 — Expected 6x6 table, found {num_rows} rows, {num_cols_row0} cols in row 0")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Row 1 headers ('', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri') (0.20 pts)
    try:
        row0_tcs = get_actual_tcs(table, 0)
        if len(row0_tcs) == 6:
            expected_headers = ['', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri']
            actual_headers = [get_tc_text(tc) for tc in row0_tcs]
            if actual_headers == expected_headers:
                print(f"PASS: Component 2 — Row 1 headers correct: {actual_headers} (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 2 — Row 1 headers mismatch. Expected: {expected_headers}, found: {actual_headers}")
        else:
            print(f"FAIL: Component 2 — Row 0 has {len(row0_tcs)} cells, expected 6")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: A2-A3 vertically merged with 'Morning' (0.20 pts)
    # A2 = row index 1, A3 = row index 2, column 0
    # Correct OOXML: row1/tc[0] has vMerge=restart + text='Morning'
    #                row2/tc[0] has vMerge=continue (empty element, no val attr)
    try:
        row1_tcs = get_actual_tcs(table, 1)
        row2_tcs = get_actual_tcs(table, 2)

        if len(row1_tcs) >= 1 and len(row2_tcs) >= 1:
            tc_a2 = row1_tcs[0]
            tc_a3 = row2_tcs[0]

            vmerge_a2 = get_tc_vmerge(tc_a2)
            vmerge_a3 = get_tc_vmerge(tc_a3)
            text_a2 = get_tc_text(tc_a2)

            if vmerge_a2 == 'restart' and vmerge_a3 == 'continue' and text_a2 == 'Morning':
                print(f"PASS: Component 3 — A2-A3 vertically merged with 'Morning' (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 3 — A2 vMerge={vmerge_a2!r} text={text_a2!r}, A3 vMerge={vmerge_a3!r}. "
                      f"Expected: A2 vMerge=restart+text='Morning', A3 vMerge=continue")
        else:
            print(f"FAIL: Component 3 — Not enough tc elements in rows 1 or 2")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: A4-A5 vertically merged with 'Afternoon' (0.20 pts)
    # A4 = row index 3, A5 = row index 4, column 0
    try:
        row3_tcs = get_actual_tcs(table, 3)
        row4_tcs = get_actual_tcs(table, 4)

        if len(row3_tcs) >= 1 and len(row4_tcs) >= 1:
            tc_a4 = row3_tcs[0]
            tc_a5 = row4_tcs[0]

            vmerge_a4 = get_tc_vmerge(tc_a4)
            vmerge_a5 = get_tc_vmerge(tc_a5)
            text_a4 = get_tc_text(tc_a4)

            if vmerge_a4 == 'restart' and vmerge_a5 == 'continue' and text_a4 == 'Afternoon':
                print(f"PASS: Component 4 — A4-A5 vertically merged with 'Afternoon' (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 4 — A4 vMerge={vmerge_a4!r} text={text_a4!r}, A5 vMerge={vmerge_a5!r}. "
                      f"Expected: A4 vMerge=restart+text='Afternoon', A5 vMerge=continue")
        else:
            print(f"FAIL: Component 4 — Not enough tc elements in rows 3 or 4")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Row 6 fully merged (gridSpan=6) with text 'Lunch' (0.20 pts)
    # Row 6 = row index 5; should have exactly 1 tc element with gridSpan=6 and text='Lunch'
    try:
        row5_tcs = get_actual_tcs(table, 5)

        if len(row5_tcs) == 1:
            tc_row6 = row5_tcs[0]
            span = get_tc_gridspan(tc_row6)
            text_r6 = get_tc_text(tc_row6)

            if span == 6 and text_r6 == 'Lunch':
                print(f"PASS: Component 5 — Row 6 merged (gridSpan=6) containing 'Lunch' (0.20 pts)")
                total_score += 0.20
            else:
                print(f"FAIL: Component 5 — Row 6 tc[0] gridSpan={span}, text={text_r6!r}. "
                      f"Expected gridSpan=6 and text='Lunch'")
        else:
            print(f"FAIL: Component 5 — Row 6 has {len(row5_tcs)} actual tc elements, expected 1 (horizontal merge)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
