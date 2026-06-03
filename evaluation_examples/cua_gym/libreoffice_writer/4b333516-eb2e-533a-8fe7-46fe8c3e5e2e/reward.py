"""
Reward Script: Configure outline numbering for Heading 1 (Chapter I, II...) and Heading 2 (1.1, 1.2...)
Task ID: writer_fp_009
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Level 1 numbering definition uses upperRoman with 'Chapter' prefix
  Component 2 (0.3): Level 2 numbering definition uses decimal with parent numbering (%1.%2)
  Component 3 (0.2): All 4 Heading 1 paragraphs are linked to the multilevel numbering
  Component 4 (0.2): All 12 Heading 2 paragraphs are linked to the multilevel numbering
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fp_009'


def persist_app_state(domain):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        import time
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for %s" % domain)
    except Exception as e:
        print("PERSIST_WARN: save hook failed: %s" % e)


def find_heading_numbering_def(doc):
    """
    Find the multilevel abstract numbering definition that is linked to
    Heading1/Heading2 styles. Returns a dict mapping ilvl -> level element,
    or empty dict if not found.
    """
    wval = qn('w:val')
    try:
        numbering_elem = doc.part.numbering_part.element
    except Exception:
        return {}, None

    # Search all abstractNum elements for a multilevel definition
    # that has pStyle linked to Heading1 or Heading2
    for abstract_num in numbering_elem.findall(qn('w:abstractNum')):
        multi_type = abstract_num.find(qn('w:multiLevelType'))
        if multi_type is not None and multi_type.get(wval) == 'multilevel':
            levels = {}
            has_heading_link = False
            for lvl in abstract_num.findall(qn('w:lvl')):
                ilvl = lvl.get(qn('w:ilvl'))
                levels[ilvl] = lvl
                pStyle = lvl.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(wval) in ('Heading1', 'Heading 1'):
                    has_heading_link = (pStyle is not None)  # derived from check
            if has_heading_link and len(levels) >= 2:
                abstract_id = abstract_num.get(qn('w:abstractNumId'))
                # Find the numId that references this abstractNum
                num_id = None
                for num_elem in numbering_elem.findall(qn('w:num')):
                    abs_ref = num_elem.find(qn('w:abstractNumId'))
                    if abs_ref is not None and abs_ref.get(wval) == abstract_id:
                        num_id = num_elem.get(qn('w:numId'))
                        break
                return levels, num_id

    # Fallback: search for any multilevel numbering with >=2 levels
    for abstract_num in numbering_elem.findall(qn('w:abstractNum')):
        multi_type = abstract_num.find(qn('w:multiLevelType'))
        if multi_type is not None and multi_type.get(wval) == 'multilevel':
            levels = {}
            for lvl in abstract_num.findall(qn('w:lvl')):
                ilvl = lvl.get(qn('w:ilvl'))
                levels[ilvl] = lvl
            if len(levels) >= 2:
                abstract_id = abstract_num.get(qn('w:abstractNumId'))
                num_id = None
                for num_elem in numbering_elem.findall(qn('w:num')):
                    abs_ref = num_elem.find(qn('w:abstractNumId'))
                    if abs_ref is not None and abs_ref.get(wval) == abstract_id:
                        num_id = num_elem.get(qn('w:numId'))
                        break
                return levels, num_id

    return {}, None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0
    wval = qn('w:val')

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    levels, num_id = find_heading_numbering_def(doc)

    # Component 1: Level 1 (ilvl=0) uses upperRoman with 'Chapter' prefix (0.3 points)
    try:
        if '0' in levels:
            lvl0 = levels['0']
            num_fmt = lvl0.find(qn('w:numFmt'))
            lvl_text = lvl0.find(qn('w:lvlText'))

            fmt_val = num_fmt.get(wval) if num_fmt is not None else None
            text_val = lvl_text.get(wval) if lvl_text is not None else None

            is_roman = fmt_val == 'upperRoman'
            has_chapter = text_val is not None and 'Chapter' in text_val and '%1' in text_val

            if is_roman and has_chapter:
                print("PASS: Component 1 -- Level 1 uses upperRoman with 'Chapter' prefix (lvlText='%s') (0.3 pts)" % text_val)
                total_score += 0.3
            elif is_roman:
                print("PARTIAL: Component 1 -- Level 1 uses upperRoman but lvlText='%s' missing 'Chapter' prefix (0.15 pts)" % text_val)
                total_score += 0.15
            elif has_chapter:
                print("PARTIAL: Component 1 -- Level 1 has 'Chapter' prefix but numFmt='%s' not upperRoman (0.15 pts)" % fmt_val)
                total_score += 0.15
            else:
                print("FAIL: Component 1 -- Level 1 numFmt='%s', lvlText='%s', expected upperRoman with 'Chapter' prefix" % (fmt_val, text_val))
        else:
            print("FAIL: Component 1 -- No level 0 found in multilevel numbering definition")
    except Exception as e:
        print("ERROR: Component 1 -- %s" % e)

    # Component 2: Level 2 (ilvl=1) uses decimal with parent numbering %1.%2 (0.3 points)
    try:
        if '1' in levels:
            lvl1 = levels['1']
            num_fmt = lvl1.find(qn('w:numFmt'))
            lvl_text = lvl1.find(qn('w:lvlText'))

            fmt_val = num_fmt.get(wval) if num_fmt is not None else None
            text_val = lvl_text.get(wval) if lvl_text is not None else None

            is_decimal = fmt_val == 'decimal'
            # Accept patterns like "%1.%2", "%1.%2 ", "%1.%2.", etc.
            has_parent_num = text_val is not None and '%1' in text_val and '%2' in text_val

            if is_decimal and has_parent_num:
                print("PASS: Component 2 -- Level 2 uses decimal with parent numbering (lvlText='%s') (0.3 pts)" % text_val)
                total_score += 0.3
            elif is_decimal:
                print("PARTIAL: Component 2 -- Level 2 uses decimal but lvlText='%s' missing parent numbering (0.15 pts)" % text_val)
                total_score += 0.15
            elif has_parent_num:
                print("PARTIAL: Component 2 -- Level 2 has parent numbering but numFmt='%s' not decimal (0.15 pts)" % fmt_val)
                total_score += 0.15
            else:
                print("FAIL: Component 2 -- Level 2 numFmt='%s', lvlText='%s', expected decimal with parent numbering" % (fmt_val, text_val))
        else:
            print("FAIL: Component 2 -- No level 1 found in multilevel numbering definition")
    except Exception as e:
        print("ERROR: Component 2 -- %s" % e)

    # Component 3: All 4 Heading 1 paragraphs are linked to the multilevel numbering (0.2 points)
    try:
        h1_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 1']
        h1_count = len(h1_paras)
        h1_linked = 0

        for p in h1_paras:
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_elem = numPr.find(qn('w:numId'))
                    ilvl_elem = numPr.find(qn('w:ilvl'))
                    nid = numId_elem.get(wval) if numId_elem is not None else None
                    ilv = ilvl_elem.get(wval) if ilvl_elem is not None else None
                    # Must reference the multilevel numbering at ilvl 0
                    if nid is not None and nid != '0' and ilv == '0':
                        h1_linked += 1

        if h1_count >= 4 and h1_linked >= 4:
            print("PASS: Component 3 -- All %d Heading 1 paragraphs linked to numbering (0.2 pts)" % h1_linked)
            total_score += 0.2
        elif h1_linked > 0:
            partial = round(0.2 * (h1_linked / max(h1_count, 4)), 2)
            if partial > 0:
                print("PARTIAL: Component 3 -- %d/%d Heading 1 paragraphs linked (%s pts)" % (h1_linked, h1_count, partial))
                total_score += partial
        else:
            print("FAIL: Component 3 -- No Heading 1 paragraphs linked to numbering (found %d Heading 1 paras)" % h1_count)
    except Exception as e:
        print("ERROR: Component 3 -- %s" % e)

    # Component 4: All 12 Heading 2 paragraphs are linked to the multilevel numbering (0.2 points)
    try:
        h2_paras = [p for p in doc.paragraphs if p.style and p.style.name == 'Heading 2']
        h2_count = len(h2_paras)
        h2_linked = 0

        for p in h2_paras:
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_elem = numPr.find(qn('w:numId'))
                    ilvl_elem = numPr.find(qn('w:ilvl'))
                    nid = numId_elem.get(wval) if numId_elem is not None else None
                    ilv = ilvl_elem.get(wval) if ilvl_elem is not None else None
                    # Must reference the multilevel numbering at ilvl 1
                    if nid is not None and nid != '0' and ilv == '1':
                        h2_linked += 1

        if h2_count >= 12 and h2_linked >= 12:
            print("PASS: Component 4 -- All %d Heading 2 paragraphs linked to numbering (0.2 pts)" % h2_linked)
            total_score += 0.2
        elif h2_linked > 0:
            partial = round(0.2 * (h2_linked / max(h2_count, 12)), 2)
            if partial > 0:
                print("PARTIAL: Component 4 -- %d/%d Heading 2 paragraphs linked (%s pts)" % (h2_linked, h2_count, partial))
                total_score += partial
        else:
            print("FAIL: Component 4 -- No Heading 2 paragraphs linked to numbering (found %d Heading 2 paras)" % h2_count)
    except Exception as e:
        print("ERROR: Component 4 -- %s" % e)

    final_score = min(total_score, 1.0)
    print("\nScore: %s/1.0" % total_score)
    print("REWARD: %s" % final_score)
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = '%s/%s.docx' % (WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
