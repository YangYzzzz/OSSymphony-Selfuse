"""
Reward Script: Appendix A page numbering with 'A-1', 'A-2' format
Task ID: writer_acad_071
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Footer contains 'A-' prefix text before page number
  Component 2 (0.3): Page numbering restarts at 1 (pgNumType start='1')
  Component 3 (0.3): Footer has both 'A-' prefix AND PAGE field (complete format)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_071'


def persist_app_state(domain: str):
    """Best-effort save any unsaved GUI edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires Appendix A pages to display 'A-1', 'A-2', etc.
    This means:
    - The appendix section footer must contain 'A-' prefix text
    - A PAGE field must follow the prefix
    - Page numbering must restart at 1 for the appendix section
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least 3 sections
    if len(doc.sections) < 3:
        print(f"FAIL: Expected at least 3 sections, found {len(doc.sections)}")
        print("REWARD: 0.0")
        return 0.0

    # The appendix section is section 2 (0-indexed)
    appendix_section = doc.sections[2]
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: Footer contains 'A-' prefix text (0.4 points)
    # In initial_env, footer just has a PAGE field with no prefix.
    # In golden_env, footer has 'A-' text followed by PAGE field.
    try:
        footer = appendix_section.footer
        has_a_prefix = False

        for para in footer.paragraphs:
            # Check runs for 'A-' text
            for run in para.runs:
                if 'A-' in run.text:
                    has_a_prefix = True
                    break
            # Also check raw XML w:t elements
            if not has_a_prefix:
                t_elements = para._element.findall('.//w:t', ns)
                for t_el in t_elements:
                    if t_el.text and 'A-' in t_el.text:
                        has_a_prefix = True
                        break

        if has_a_prefix:
            print(f"PASS: Component 1 - Footer contains 'A-' prefix text (0.4 pts)")
            total_score += 0.4
        else:
            # Gather what footer actually shows
            footer_texts = [p.text for p in footer.paragraphs]
            print(f"FAIL: Component 1 - Footer does not contain 'A-' prefix. Footer text: {footer_texts}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Page numbering restarts at 1 (pgNumType start='1') (0.3 points)
    # In initial_env, section 2 has pgNumType with fmt='decimal' but no start attribute.
    # In golden_env, section 2 has pgNumType with start='1'.
    try:
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        sectPr = appendix_section._sectPr
        pgNumType = sectPr.find('{%s}pgNumType' % ns_w)

        if pgNumType is not None:
            start_val = pgNumType.get('{%s}start' % ns_w)
            if start_val == '1':
                print(f"PASS: Component 2 - Page numbering restarts at 1 (pgNumType start=1) (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 - pgNumType found but start={start_val!r}, expected '1'")
        else:
            print(f"FAIL: Component 2 - No pgNumType element in appendix section")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Footer has BOTH 'A-' prefix AND PAGE field code (0.3 points)
    # This compound check ensures the complete 'A-N' format is present.
    # In initial_env: has PAGE field but no 'A-' prefix -> FAIL
    # In golden_env: has both 'A-' prefix and PAGE field -> PASS
    try:
        footer = appendix_section.footer
        has_a_prefix_c3 = False
        has_page_field = False

        for para in footer.paragraphs:
            # Check for 'A-' prefix
            t_elements = para._element.findall('.//w:t', ns)
            for t_el in t_elements:
                if t_el.text and 'A-' in t_el.text:
                    has_a_prefix_c3 = True

            # Check for PAGE field code
            instr_texts = para._element.findall('.//w:instrText', ns)
            for instr in instr_texts:
                if instr.text and 'PAGE' in instr.text:
                    has_page_field = True

        if has_a_prefix_c3 and has_page_field:
            print(f"PASS: Component 3 - Footer has 'A-' prefix AND PAGE field (complete A-N format) (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 - Incomplete format: A-prefix={has_a_prefix_c3}, PAGE-field={has_page_field}")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
