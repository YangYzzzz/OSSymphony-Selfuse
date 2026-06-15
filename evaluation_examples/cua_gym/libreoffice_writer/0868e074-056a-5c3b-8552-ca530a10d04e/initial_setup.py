"""
Initial Setup: Thesis document with Appendix A using continuous Arabic numbering
Task ID: writer_acad_071
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_071'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section, prefix_text=""):
    """Add a footer with optional prefix text and a PAGE field code."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if prefix_text:
        run_prefix = fp.add_run(prefix_text)
        run_prefix.font.size = Pt(10)
        run_prefix.font.name = "Times New Roman"

    # PAGE field code
    r1 = fp.add_run()
    r1.font.size = Pt(10)
    r1.font.name = "Times New Roman"
    fldChar1 = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fldChar1)

    r2 = fp.add_run()
    r2.font.size = Pt(10)
    r2.font.name = "Times New Roman"
    instrText = r2._element.makeelement(qn('w:instrText'), {})
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' PAGE '
    r2._element.append(instrText)

    r3 = fp.add_run()
    r3.font.size = Pt(10)
    r3.font.name = "Times New Roman"
    fldChar3 = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fldChar3)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ====== MAIN BODY ======

    # --- Title Page ---
    section0 = doc.sections[0]
    section0.top_margin = Inches(2)
    section0.left_margin = Inches(1.25)
    section0.right_margin = Inches(1.25)
    section0.bottom_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(72)
    run_t = title.add_run("The Impact of Machine Learning Algorithms\non Urban Traffic Flow Optimization")
    run_t.bold = True
    run_t.font.size = Pt(16)
    run_t.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(48)
    run_s = subtitle.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nMaster of Science in Computer Engineering")
    run_s.font.size = Pt(12)
    run_s.font.name = "Times New Roman"

    author = doc.add_paragraph()
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(36)
    run_a = author.add_run("by\nElena Vasquez Rodriguez")
    run_a.font.size = Pt(12)
    run_a.font.name = "Times New Roman"

    dept = doc.add_paragraph()
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept.paragraph_format.space_before = Pt(24)
    run_d = dept.add_run("Department of Computer Engineering\nNorthwestern Polytechnic University\nMay 2025")
    run_d.font.size = Pt(12)
    run_d.font.name = "Times New Roman"

    # No footer on title page
    footer0 = section0.footer
    footer0.is_linked_to_previous = False

    # --- Chapter 1: Introduction (new section, new page) ---
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section1 = doc.sections[1]
    section1.left_margin = Inches(1.25)
    section1.right_margin = Inches(1.25)
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)

    # Add page number footer for body (just Arabic numbers)
    add_page_number_footer(section1)

    h1 = doc.add_heading("Chapter 1: Introduction", level=1)
    for run in h1.runs:
        run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Urban traffic congestion remains one of the most pressing challenges facing "
        "metropolitan areas worldwide. According to the Texas A&M Transportation Institute, "
        "American commuters spent an average of 54 extra hours in traffic in 2023, resulting "
        "in approximately $87 billion in lost productivity and wasted fuel. Traditional traffic "
        "management systems, relying on fixed-time signal controllers and predetermined routing "
        "algorithms, have proven inadequate in addressing the dynamic and non-linear nature of "
        "modern urban traffic patterns."
    )
    doc.add_paragraph(
        "Recent advances in machine learning, particularly in deep reinforcement learning and "
        "graph neural networks, offer promising alternatives. These approaches can model complex "
        "spatial-temporal dependencies in traffic data and adapt in real-time to changing conditions. "
        "However, significant challenges remain in deploying these models at scale, including "
        "computational constraints, data quality issues, and the need for interpretable decision-making "
        "in safety-critical transportation systems."
    )
    doc.add_paragraph(
        "This thesis investigates the application of three distinct machine learning paradigms to "
        "the problem of urban traffic flow optimization: (1) deep reinforcement learning for adaptive "
        "signal control, (2) graph convolutional networks for network-wide traffic prediction, and "
        "(3) federated learning for privacy-preserving multi-intersection coordination. Our approach "
        "is evaluated using both simulated environments and real-world data from the City of Portland "
        "Automated Traffic Signal Performance Measures (ATSPM) database."
    )

    # --- Chapter 2: Literature Review ---
    h2 = doc.add_heading("Chapter 2: Literature Review", level=1)
    for run in h2.runs:
        run.font.name = "Times New Roman"

    doc.add_paragraph(
        "The application of computational intelligence to traffic management has evolved significantly "
        "over the past two decades. Early work by Srinivasan, Choy, and Cheu (2006) demonstrated that "
        "neural networks could outperform traditional Webster timing plans for isolated intersections. "
        "Subsequent research by Abdulhai, Pringle, and Karakoulas (2003) introduced Q-learning for "
        "single-intersection signal control, achieving a 28% reduction in average vehicle delay "
        "compared to actuated control methods."
    )
    doc.add_paragraph(
        "The emergence of deep reinforcement learning (DRL) marked a paradigm shift in this field. "
        "Wei et al. (2018) proposed IntelliLight, combining deep Q-networks with phase-based state "
        "representations, while Zheng et al. (2019) introduced FRAP, which leveraged the symmetric "
        "structure of traffic phases to improve learning efficiency. More recently, multi-agent "
        "reinforcement learning approaches, such as CoLight (Wei et al., 2019) and PressLight "
        "(Wei et al., 2019), have extended DRL to network-level coordination, demonstrating the "
        "feasibility of decentralized yet coordinated signal control."
    )
    doc.add_paragraph(
        "Graph neural networks (GNNs) have also gained traction for traffic prediction tasks. "
        "Li et al. (2018) proposed Diffusion Convolutional Recurrent Neural Networks (DCRNN) for "
        "traffic speed forecasting, while Yu, Yin, and Zhu (2018) introduced Spatio-Temporal Graph "
        "Convolutional Networks (STGCN). These models capture both spatial correlations across road "
        "network topology and temporal dependencies in traffic dynamics, achieving state-of-the-art "
        "prediction accuracy on benchmark datasets including METR-LA and PEMS-BAY."
    )

    # --- Chapter 3: Methodology ---
    h3 = doc.add_heading("Chapter 3: Methodology", level=1)
    for run in h3.runs:
        run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Our experimental framework integrates three complementary machine learning approaches within "
        "a unified simulation environment built on SUMO (Simulation of Urban Mobility) version 1.15.0. "
        "The simulation network models a 4.2 km² area of downtown Portland, Oregon, comprising 47 "
        "signalized intersections connected by 132 road segments with varying lane configurations "
        "and speed limits ranging from 25 to 45 mph."
    )
    doc.add_paragraph(
        "For the deep reinforcement learning component, we implement a Proximal Policy Optimization "
        "(PPO) agent with a custom observation space encoding current phase state, queue lengths, "
        "vehicle counts per lane, and waiting time distributions. The reward function balances multiple "
        "objectives: minimizing average intersection delay (weighted at 0.4), reducing queue spillback "
        "incidents (weighted at 0.3), and maintaining equitable green time distribution across competing "
        "movements (weighted at 0.3). Training utilizes 500 episodes of 3,600-second simulations with "
        "demand patterns derived from Portland's ATSPM detector data for Q3 2024."
    )

    # --- Chapter 4: Results ---
    h4 = doc.add_heading("Chapter 4: Results and Discussion", level=1)
    for run in h4.runs:
        run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Table 4.1 summarizes the performance comparison across all evaluated approaches under "
        "three demand scenarios: low (60% of peak), medium (85% of peak), and high (100% of peak). "
        "Our integrated ML framework achieves statistically significant improvements over the baseline "
        "actuated control system across all metrics and demand levels (p < 0.001, paired t-test with "
        "Bonferroni correction)."
    )

    # Add a simple results table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Method', 'Avg Delay (s)', 'Queue Length (m)', 'Throughput (veh/h)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    data_rows = [
        ['Fixed-Time Control', '47.3 ± 4.1', '128.5 ± 15.2', '1,842 ± 98'],
        ['Actuated Control', '38.6 ± 3.8', '96.3 ± 12.7', '2,105 ± 87'],
        ['DRL (PPO)', '26.4 ± 2.9', '62.1 ± 8.4', '2,487 ± 72'],
        ['Integrated ML (Ours)', '21.8 ± 2.3', '48.7 ± 6.9', '2,651 ± 65'],
    ]
    for r, row_data in enumerate(data_rows, 1):
        for c, val in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    doc.add_paragraph(
        "The integrated framework demonstrates a 43.5% reduction in average delay compared to "
        "actuated control under high-demand conditions. Graph-based prediction enables proactive "
        "signal adjustments 3-5 minutes ahead of demand surges, while the federated coordination "
        "layer ensures consistent performance across the network without requiring centralized data "
        "aggregation."
    )

    # ====== APPENDIX A (same section, continuous numbering - this is the initial "problem") ======
    # Add a new section for Appendix A, but keep continuous numbering (the problem to fix)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_app = doc.sections[-1]
    section_app.left_margin = Inches(1.25)
    section_app.right_margin = Inches(1.25)
    section_app.top_margin = Inches(1)
    section_app.bottom_margin = Inches(1)

    # Footer linked to previous - continues Arabic numbering
    footer_app = section_app.footer
    footer_app.is_linked_to_previous = True

    h_app = doc.add_heading("Appendix A: Supplementary Data Tables", level=1)
    for run in h_app.runs:
        run.font.name = "Times New Roman"

    doc.add_paragraph(
        "This appendix contains the complete experimental data tables referenced throughout "
        "the main body of the thesis. Table A.1 provides the detailed intersection-level "
        "performance metrics for all 47 signalized intersections in the study network under "
        "high-demand conditions."
    )

    # Table A.1
    p_tbl = doc.add_paragraph()
    p_tbl.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_tbl = p_tbl.add_run("Table A.1: Intersection-Level Performance Metrics (High Demand)")
    run_tbl.bold = True
    run_tbl.font.size = Pt(10)
    run_tbl.font.name = "Times New Roman"

    tbl_a1 = doc.add_table(rows=13, cols=5)
    tbl_a1.style = 'Table Grid'
    a1_headers = ['Intersection ID', 'Avg Delay (s)', 'Max Queue (m)', 'LOS Grade', 'Volume (veh/h)']
    for i, h in enumerate(a1_headers):
        cell = tbl_a1.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

    intersection_data = [
        ['INT-001 (MLK & Broadway)', '18.4', '42.3', 'B', '1,245'],
        ['INT-002 (Burnside & 3rd)', '31.7', '78.9', 'C', '1,892'],
        ['INT-003 (Powell & 82nd)', '45.2', '112.5', 'D', '2,134'],
        ['INT-004 (Sandy & 47th)', '22.1', '53.7', 'C', '987'],
        ['INT-005 (Hawthorne & 39th)', '27.8', '64.2', 'C', '1,456'],
        ['INT-006 (Division & 60th)', '38.9', '95.1', 'D', '1,678'],
        ['INT-007 (Alberta & MLK)', '15.3', '35.8', 'B', '876'],
        ['INT-008 (Stark & 12th)', '52.6', '134.7', 'E', '2,345'],
        ['INT-009 (Lombard & Interstate)', '29.4', '71.3', 'C', '1,567'],
        ['INT-010 (Glisan & 33rd)', '24.6', '58.4', 'C', '1,123'],
        ['INT-011 (Cesar Chavez & Belmont)', '33.1', '82.6', 'C', '1,389'],
        ['INT-012 (Foster & 72nd)', '41.7', '103.2', 'D', '1,756'],
    ]
    for r, row_data in enumerate(intersection_data, 1):
        for c, val in enumerate(row_data):
            cell = tbl_a1.cell(r, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)

    doc.add_paragraph()  # spacing

    doc.add_paragraph(
        "Table A.2 presents the model hyperparameters used in the final training configuration "
        "for each of the three machine learning components. All hyperparameters were selected "
        "through Bayesian optimization using Optuna with 200 trials per component."
    )

    # Table A.2
    p_tbl2 = doc.add_paragraph()
    p_tbl2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_tbl2 = p_tbl2.add_run("Table A.2: Model Hyperparameters")
    run_tbl2.bold = True
    run_tbl2.font.size = Pt(10)
    run_tbl2.font.name = "Times New Roman"

    tbl_a2 = doc.add_table(rows=9, cols=3)
    tbl_a2.style = 'Table Grid'
    a2_headers = ['Parameter', 'Component', 'Value']
    for i, h in enumerate(a2_headers):
        cell = tbl_a2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    hyperparam_data = [
        ['Learning Rate', 'DRL (PPO)', '3.0 × 10⁻⁴'],
        ['Discount Factor (γ)', 'DRL (PPO)', '0.99'],
        ['Clip Range (ε)', 'DRL (PPO)', '0.2'],
        ['Hidden Layers', 'GCN', '[128, 64, 32]'],
        ['Dropout Rate', 'GCN', '0.15'],
        ['Graph Attention Heads', 'GCN', '4'],
        ['Local Epochs', 'Federated', '5'],
        ['Aggregation Rounds', 'Federated', '50'],
    ]
    for r, row_data in enumerate(hyperparam_data, 1):
        for c, val in enumerate(row_data):
            cell = tbl_a2.cell(r, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    doc.add_paragraph(
        "Additional calibration results and convergence plots for the reinforcement learning "
        "training process are available in the supplementary digital materials repository "
        "hosted at https://github.com/evasquez/traffic-ml-thesis."
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
