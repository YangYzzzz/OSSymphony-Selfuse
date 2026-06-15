"""
Initial Setup: Novel manuscript with chapter title (no outline effect)
Task ID: writer_txtfmt_026
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user/Desktop'
TASK_ID = 'writer_txtfmt_026'
OUTPUT = f'{WORKDIR}/novel_draft.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # --- Chapter Title: 18pt Georgia Bold, NO outline effect ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Chapter 1: The Beginning")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.outline = False  # explicitly no outline in initial state

    # --- Narrative paragraphs in 12pt Georgia ---
    paragraphs = [
        (
            "The morning mist clung to the valley like a whispered secret, "
            "reluctant to release its hold on the sleeping village below. "
            "Eleanor stood at the edge of Thornwood Cliff, her worn leather "
            "satchel hanging from one shoulder, watching the sun slowly burn "
            "away the fog."
        ),
        (
            "She had arrived in Millhaven three days ago, following a map drawn "
            "by her grandmother's trembling hand on the back of an old envelope. "
            "The village itself was unremarkable — a cluster of stone cottages, a "
            "crooked church steeple, and a marketplace that smelled of fresh bread "
            "and salt air. Yet something about the place felt deeply familiar, as "
            "though she had walked these cobblestone streets in another life."
        ),
        (
            "\"You look lost,\" said a voice behind her. Eleanor turned to find a "
            "young man in a fisherman's coat, his dark eyes curious and kind. He "
            "carried a crate of mackerel in both arms, which he set down on a "
            "nearby wall with practiced ease."
        ),
        (
            "\"Not lost,\" she replied, tucking a strand of auburn hair behind her "
            "ear. \"Just looking for something I'm not sure I'll recognize when I "
            "find it.\""
        ),
        (
            "He laughed — a warm, unhurried sound that echoed off the cliff face. "
            "\"That sounds like half the people in this town. I'm Daniel Marsh. "
            "My family has fished these waters for six generations.\""
        ),
        (
            "Eleanor introduced herself and shook his calloused hand. As she did, "
            "the leather satchel slipped from her shoulder and snapped open, "
            "spilling its contents across the damp grass: a compass, two slim "
            "notebooks bound in twine, a photograph of a woman she had never met, "
            "and the envelope with the map."
        ),
        (
            "Daniel helped her gather everything without comment, but his gaze "
            "lingered on the photograph. When he handed it back to her, his "
            "expression had changed — less open, more guarded, as if a door had "
            "quietly closed somewhere behind his eyes."
        ),
    ]

    for text in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Georgia"
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
