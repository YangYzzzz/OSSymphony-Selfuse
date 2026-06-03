"""
Reward Script: Insert a Database record number field in mail merge letter
Task ID: writer_mt_037
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30): "Letter X of 50" text present at top of document
  Component 2 (0.20): The "Letter ... of 50" paragraph is right-aligned
  Component 3 (0.30): MERGEREC field code present (database record number field)
  Component 4 (0.10): "of 50" is plain text (manually typed)
  Component 5 (0.10): Original document content preserved below
"""

import os
import re
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_037'
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # We need at least 2 paragraphs (the new "Letter X of 50" + original content)
    if len(doc.paragraphs) < 2:
        print("FAIL: Document has fewer than 2 paragraphs")
        print("REWARD: 0.0")
        return 0.0

    # Search through all paragraphs for one matching "Letter ... of 50"
    letter_para = None
    letter_para_idx = None
    pattern = re.compile(r'Letter\s+\d+\s+of\s+50', re.IGNORECASE)
    for i, para in enumerate(doc.paragraphs):
        if pattern.search(para.text):
            letter_para = para
            letter_para_idx = i
            break

    # Component 1: "Letter X of 50" text present (0.30 points)
    try:
        if letter_para is not None:
            print(f"PASS: Component 1 -- 'Letter X of 50' text found in para {letter_para_idx}: {letter_para.text!r} (0.30 pts)")
            total_score += 0.30
        else:
            # Also check headers for this text
            found_in_header = False
            for section in doc.sections:
                for p in section.header.paragraphs:
                    if pattern.search(p.text):
                        letter_para = p
                        found_in_header = True
                        print(f"PASS: Component 1 -- 'Letter X of 50' text found in header: {p.text!r} (0.30 pts)")
                        total_score += 0.30
                        break
                if found_in_header:
                    break
            if not found_in_header:
                all_texts = [p.text for p in doc.paragraphs[:5]]
                print(f"FAIL: Component 1 -- No 'Letter X of 50' text found. First 5 paras: {all_texts}")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    if letter_para is None:
        # Cannot proceed with further checks without the target paragraph
        final_score = min(total_score, 1.0)
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {final_score}")
        return final_score

    # Component 2: Right-aligned paragraph (0.20 points)
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        alignment = letter_para.paragraph_format.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            print(f"PASS: Component 2 -- Paragraph is right-aligned (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 -- Expected RIGHT alignment, found: {alignment}")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: MERGEREC field code present in the paragraph (0.30 points)
    # The record number should be a database field, not plain text
    try:
        para_xml = etree.tostring(letter_para._element, pretty_print=True).decode()
        # Check for MERGEREC or DATABASE or MERGEFIELD field codes related to record number
        instr_texts = letter_para._element.findall('.//w:instrText', NS)
        field_found = False
        for instr in instr_texts:
            instr_val = (instr.text or '').strip().upper()
            # MERGEREC is the standard field for record number in mail merge
            # Also accept DATABASE, MERGESEQ or similar record number fields
            if any(kw in instr_val for kw in ['MERGEREC', 'DATABASE', 'MERGESEQ']):
                field_found = True
                print(f"PASS: Component 3 -- Database/merge record field found: {instr_val!r} (0.30 pts)")
                total_score += 0.30
                break

        if not field_found:
            # Also check for fldChar elements as indicator of any field
            fld_chars = letter_para._element.findall('.//w:fldChar', NS)
            if fld_chars:
                # There are field codes but not the expected ones
                all_instr = [i.text for i in instr_texts]
                print(f"FAIL: Component 3 -- Field codes found but not MERGEREC/DATABASE: {all_instr}")
            else:
                print(f"FAIL: Component 3 -- No field codes found in the paragraph (record number appears to be plain text)")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: "of 50" is plain text, not a field (0.10 points)
    try:
        # Check that "of 50" appears as regular text (in a run that is NOT inside a field)
        runs = letter_para.runs
        of50_found = False
        for run in runs:
            if 'of 50' in run.text:
                of50_found = True
                break
        if of50_found:
            print(f"PASS: Component 4 -- 'of 50' is plain text in a run (0.10 pts)")
            total_score += 0.10
        else:
            # Check para.text as fallback (field display values may merge)
            if 'of 50' in letter_para.text:
                print(f"PASS: Component 4 -- 'of 50' found in paragraph text (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 4 -- 'of 50' not found in paragraph text: {letter_para.text!r}")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    # Component 5: Original document content preserved (0.10 points)
    # The original document starts with "Meridian Financial Services"
    # Check that this text still exists somewhere in the document
    try:
        all_text = ' '.join(p.text for p in doc.paragraphs)
        checks_passed = 0
        checks_total = 3

        if 'Meridian Financial Services' in all_text:
            checks_passed += 1
        else:
            print(f"FAIL: Component 5 -- 'Meridian Financial Services' not found in document")

        if 'Dear' in all_text and 'FirstName' in all_text:
            checks_passed += 1
        else:
            print(f"FAIL: Component 5 -- Merge greeting 'Dear FirstName' not found")

        if 'Sincerely' in all_text:
            checks_passed += 1
        else:
            print(f"FAIL: Component 5 -- 'Sincerely' closing not found")

        if checks_passed == checks_total:
            print(f"PASS: Component 5 -- Original document content preserved ({checks_passed}/{checks_total} checks) (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 5 -- Only {checks_passed}/{checks_total} content checks passed")
    except Exception as e:
        print(f"ERROR: Component 5 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
