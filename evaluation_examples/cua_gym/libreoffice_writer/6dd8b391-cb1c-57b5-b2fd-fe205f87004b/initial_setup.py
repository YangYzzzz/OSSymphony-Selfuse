"""
Initial Setup: Mail merge letter template without record number field
Task ID: writer_mt_037
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_037'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_merge_field(paragraph, field_name):
    """Add a MERGEFIELD field code to a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = run2._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = f' MERGEFIELD {field_name} '
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_separate = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    run3._element.append(fld_char_separate)

    run4 = paragraph.add_run(f'\u00ab{field_name}\u00bb')
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    run5 = paragraph.add_run()
    fld_char_end = run5._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run5._element.append(fld_char_end)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # -- Company header --
    header_para = doc.add_paragraph()
    header_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.paragraph_format.space_after = Pt(6)
    run = header_para.add_run("Meridian Financial Services")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_header = doc.add_paragraph()
    sub_header.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_header.paragraph_format.space_after = Pt(2)
    run = sub_header.add_run("1200 Harbor Boulevard, Suite 400  |  San Diego, CA 92101")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    sub_header2 = doc.add_paragraph()
    sub_header2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_header2.paragraph_format.space_after = Pt(12)
    run = sub_header2.add_run("Phone: (619) 555-0142  |  Fax: (619) 555-0143  |  www.meridianfs.com")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # -- Horizontal rule (thin line) --
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(0)
    hr_para.paragraph_format.space_after = Pt(12)
    pBdr = hr_para._element.get_or_add_pPr().makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '1F497D',
    })
    pBdr.append(bottom)
    hr_para._element.get_or_add_pPr().append(pBdr)

    # -- Date --
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    date_para.paragraph_format.space_after = Pt(12)
    run = date_para.add_run("March 15, 2026")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # -- Recipient address with merge fields --
    addr_para = doc.add_paragraph()
    addr_para.paragraph_format.space_after = Pt(0)
    add_merge_field(addr_para, "FirstName")
    addr_para.add_run(" ").font.size = Pt(11)
    add_merge_field(addr_para, "LastName")

    addr2 = doc.add_paragraph()
    addr2.paragraph_format.space_after = Pt(0)
    add_merge_field(addr2, "Address")

    addr3 = doc.add_paragraph()
    addr3.paragraph_format.space_after = Pt(12)
    add_merge_field(addr3, "City")
    addr3.add_run(", ").font.size = Pt(11)
    add_merge_field(addr3, "State")
    addr3.add_run(" ").font.size = Pt(11)
    add_merge_field(addr3, "ZipCode")

    # -- Salutation --
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(12)
    salutation.add_run("Dear ").font.size = Pt(11)
    add_merge_field(salutation, "FirstName")
    salutation.add_run(",").font.size = Pt(11)

    # -- Body paragraphs --
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(10)
    body1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body1.add_run(
        "Thank you for choosing Meridian Financial Services as your investment partner. "
        "We are writing to provide you with an important update regarding your portfolio "
        "performance for the first quarter of 2026."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(10)
    body2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body2.add_run(
        "As of February 28, 2026, your account balance stands at "
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    add_merge_field(body2, "AccountBalance")
    run2 = body2.add_run(
        ". This reflects a net growth of "
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"
    add_merge_field(body2, "GrowthPercent")
    run3 = body2.add_run(
        " over the past quarter, outperforming the benchmark index by 1.3 percentage points."
    )
    run3.font.size = Pt(11)
    run3.font.name = "Calibri"

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(10)
    body3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body3.add_run(
        "We recommend scheduling a review meeting with your dedicated advisor, "
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    add_merge_field(body3, "AdvisorName")
    run2 = body3.add_run(
        ", to discuss potential adjustments to your investment strategy for the coming quarter. "
        "Please contact our office at your earliest convenience to arrange an appointment."
    )
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    body4 = doc.add_paragraph()
    body4.paragraph_format.space_after = Pt(10)
    body4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = body4.add_run(
        "Enclosed you will find your detailed quarterly statement, including a breakdown "
        "of all transactions, dividend payments, and fee assessments for the period. "
        "Please review the enclosed documents carefully and do not hesitate to reach out "
        "with any questions."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # -- Closing --
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(4)
    run = closing.add_run("Sincerely,")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(24)
    sig.paragraph_format.space_after = Pt(0)
    run = sig.add_run("Robert A. Whitfield")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(0)
    run = title_para.add_run("Senior Vice President, Client Relations")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    dept_para = doc.add_paragraph()
    dept_para.paragraph_format.space_after = Pt(0)
    run = dept_para.add_run("Meridian Financial Services")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # -- Footer with confidentiality notice --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = fp.add_run("CONFIDENTIAL - This letter contains privileged financial information intended solely for the named recipient.")
    run.font.size = Pt(7)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
