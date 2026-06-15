"""
Reward Script: Create return address labels using Avery 5167 format
Task ID: writer_mt_050
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Table grid is 20 rows x 4 cols (80 labels)
  Component 2 (0.3): All cells contain correct address text
  Component 3 (0.2): First line 'Acme Corp' is bold in all cells
  Component 4 (0.2): Font size is 7pt across all cells
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_050'

EXPECTED_TEXT = 'Acme Corp\n456 Oak Ave\nPortland, OR 97201'
EXPECTED_ROWS = 20
EXPECTED_COLS = 4
EXPECTED_LABELS = EXPECTED_ROWS * EXPECTED_COLS  # 80


def persist_app_state(domain):
    """Best-effort save of any unsaved GUI edits."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for %s" % domain)
    except Exception as e:
        print("PERSIST_WARN: save hook failed: %s" % e)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
    except ImportError:
        print("CRITICAL: python-docx not installed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Table exists with 20 rows x 4 cols (0.3 points)
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 1 -- No tables found in document")
        else:
            t = doc.tables[0]
            nrows = len(t.rows)
            ncols = len(t.columns)
            if nrows == EXPECTED_ROWS and ncols == EXPECTED_COLS:
                print("PASS: Component 1 -- Table is %dx%d = %d labels (0.3 pts)" % (nrows, ncols, nrows * ncols))
                total_score += 0.3
            elif nrows * ncols >= 60:
                # Partial: at least 60 labels
                partial = 0.15
                print("PARTIAL: Component 1 -- Table is %dx%d = %d labels, expected %dx%d (%.2f pts)" % (nrows, ncols, nrows * ncols, EXPECTED_ROWS, EXPECTED_COLS, partial))
                total_score += partial
            else:
                print("FAIL: Component 1 -- Table is %dx%d = %d labels, expected %dx%d" % (nrows, ncols, nrows * ncols, EXPECTED_ROWS, EXPECTED_COLS))
    except Exception as e:
        print("ERROR: Component 1 -- %s" % e)

    # Component 2: All cells contain correct address text (0.3 points)
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 2 -- No table to check")
        else:
            t = doc.tables[0]
            correct_count = 0
            total_cells = 0
            for row in t.rows:
                for cell in row.cells:
                    total_cells += 1
                    cell_text = cell.text.strip()
                    if cell_text == EXPECTED_TEXT.strip():
                        correct_count += 1
            if total_cells > 0 and correct_count == total_cells:
                print("PASS: Component 2 -- All %d cells have correct address text (0.3 pts)" % total_cells)
                total_score += 0.3
            elif correct_count > 0:
                ratio = correct_count / total_cells
                partial = round(0.3 * ratio, 2)
                print("PARTIAL: Component 2 -- %d/%d cells correct (%.2f pts)" % (correct_count, total_cells, partial))
                total_score += partial
            else:
                print("FAIL: Component 2 -- 0/%d cells have correct text" % total_cells)
    except Exception as e:
        print("ERROR: Component 2 -- %s" % e)

    # Component 3: First line 'Acme Corp' is bold (0.2 points)
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 3 -- No table to check")
        else:
            t = doc.tables[0]
            bold_count = 0
            checked = 0
            for row in t.rows:
                for cell in row.cells:
                    if len(cell.paragraphs) >= 1:
                        first_para = cell.paragraphs[0]
                        if first_para.text.strip() == 'Acme Corp':
                            checked += 1
                            # Check if any run in first paragraph is bold
                            has_bold = any(r.font.bold is True for r in first_para.runs if r.text.strip())
                            if has_bold:
                                bold_count += 1
            if checked > 0 and bold_count == checked:
                print("PASS: Component 3 -- 'Acme Corp' is bold in all %d labels (0.2 pts)" % bold_count)
                total_score += 0.2
            elif bold_count > 0:
                ratio = bold_count / max(checked, 1)
                partial = round(0.2 * ratio, 2)
                print("PARTIAL: Component 3 -- Bold in %d/%d labels (%.2f pts)" % (bold_count, checked, partial))
                total_score += partial
            else:
                print("FAIL: Component 3 -- 'Acme Corp' is not bold (0/%d checked)" % checked)
    except Exception as e:
        print("ERROR: Component 3 -- %s" % e)

    # Component 4: Font size is 7pt across cells (0.2 points)
    try:
        if len(doc.tables) < 1:
            print("FAIL: Component 4 -- No table to check")
        else:
            t = doc.tables[0]
            cells_with_7pt = 0
            cells_checked = 0
            for row in t.rows:
                for cell in row.cells:
                    if not cell.text.strip():
                        continue
                    cells_checked += 1
                    all_runs_7pt = True
                    has_runs = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                            has_runs = True
                            if run.font.size is None or abs(run.font.size.pt - 7.0) > 0.5:
                                all_runs_7pt = False
                    if has_runs and all_runs_7pt:
                        cells_with_7pt += 1
            if cells_checked > 0 and cells_with_7pt == cells_checked:
                print("PASS: Component 4 -- All %d cells have 7pt font (0.2 pts)" % cells_with_7pt)
                total_score += 0.2
            elif cells_with_7pt > 0:
                ratio = cells_with_7pt / max(cells_checked, 1)
                partial = round(0.2 * ratio, 2)
                print("PARTIAL: Component 4 -- %d/%d cells at 7pt (%.2f pts)" % (cells_with_7pt, cells_checked, partial))
                total_score += partial
            else:
                print("FAIL: Component 4 -- No cells have 7pt font size")
    except Exception as e:
        print("ERROR: Component 4 -- %s" % e)

    final_score = min(total_score, 1.0)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = os.path.join(WORKDIR, TASK_ID + '.docx')
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
