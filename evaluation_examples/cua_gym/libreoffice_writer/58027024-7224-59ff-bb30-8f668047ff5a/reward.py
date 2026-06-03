"""
Reward Script: Leave policy table creation in Writer document
Task ID: writer_hr_030
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25) - Table exists with 6 rows x 4 columns
  Component 2 (0.25) - Header row has correct column names
  Component 3 (0.30) - Data rows contain correct leave type info
  Component 4 (0.20) - Header row has light blue background shading
"""

import os
from math import sqrt

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_030'


def is_light_blue(hex_color):
    """Check if a hex color string is perceptually close to light blue."""
    if not hex_color or hex_color.lower() in ('auto', 'none', ''):
        return False
    try:
        hex_color = hex_color.lstrip('#')
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        # Light blue: high blue, moderate-high green, moderate red
        # Must be clearly blue-ish and light (not dark blue or white)
        # BDD7EE is the expected value (r=189, g=215, b=238)
        # Accept a range of light blues
        if b < 180:
            return False
        if b < g:
            return False
        if r > b:
            return False
        # Ensure it's not just white/gray
        if abs(r - g) < 10 and abs(g - b) < 10:
            return False
        return True
    except (ValueError, IndexError):
        return False


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    from docx import Document
    from docx.oxml.ns import qn

    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in the document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Component 1: Table has correct dimensions - 6 rows x 4 columns (0.25 points)
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows == 6 and num_cols == 4:
            print(f"PASS: Component 1 - Table has {num_rows} rows x {num_cols} cols (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 - Expected 6x4 table, found {num_rows}x{num_cols}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Header row has correct column names (0.25 points)
    try:
        expected_headers = ['Leave Type', 'Days Per Year', 'Eligibility', 'Carryover Policy']
        actual_headers = [table.rows[0].cells[ci].text.strip() for ci in range(min(len(table.columns), 4))]
        matches = sum(1 for exp, act in zip(expected_headers, actual_headers) if exp.lower() == act.lower())
        if matches == 4:
            print(f"PASS: Component 2 - All 4 header columns match (0.25 pts)")
            total_score += 0.25
        elif matches >= 2:
            partial = round(0.25 * matches / 4, 2)
            print(f"PARTIAL: Component 2 - {matches}/4 headers match ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - Headers: {actual_headers}, expected: {expected_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Data rows contain correct leave type information (0.30 points)
    # Check leave types and their days/duration values
    try:
        expected_data = {
            'Annual Leave': '15',
            'Sick Leave': '10',
            'Personal Leave': '3',
            'Bereavement Leave': '5',
            'Parental Leave': '12',
        }
        found_count = 0
        for ri in range(1, min(len(table.rows), 7)):
            leave_type = table.rows[ri].cells[0].text.strip()
            days_text = table.rows[ri].cells[1].text.strip()
            for exp_type, exp_num in expected_data.items():
                if exp_type.lower() in leave_type.lower():
                    if exp_num in days_text:
                        found_count += 1
                        break
        if found_count == 5:
            print(f"PASS: Component 3 - All 5 leave types with correct values found (0.30 pts)")
            total_score += 0.30
        elif found_count >= 1:
            partial = round(0.30 * found_count / 5, 2)
            print(f"PARTIAL: Component 3 - {found_count}/5 leave types correct ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 - No matching leave type data found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Header row has light blue background shading (0.20 points)
    try:
        blue_cells = 0
        for ci in range(min(len(table.columns), 4)):
            cell = table.rows[0].cells[ci]
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if is_light_blue(fill):
                        blue_cells += 1
        if blue_cells == 4:
            print(f"PASS: Component 4 - All 4 header cells have light blue shading (0.20 pts)")
            total_score += 0.20
        elif blue_cells >= 1:
            partial = round(0.20 * blue_cells / 4, 2)
            print(f"PARTIAL: Component 4 - {blue_cells}/4 header cells have light blue shading ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - No header cells have light blue background")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
