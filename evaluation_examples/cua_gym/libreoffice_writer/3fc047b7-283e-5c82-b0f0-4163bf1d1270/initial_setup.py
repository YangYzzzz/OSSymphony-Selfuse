"""
Initial Setup: Shopping list with default round bullets
Task ID: writer_lec_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Add a title
    heading = doc.add_heading("Shopping List", level=1)

    # Add 8 grocery items as a bulleted list using the default "List Bullet" style
    grocery_items = [
        "Whole milk (1 gallon)",
        "Organic brown eggs (dozen)",
        "Sourdough bread loaf",
        "Fresh Atlantic salmon fillet",
        "Baby spinach (5 oz bag)",
        "Greek yogurt (vanilla, 32 oz)",
        "Extra virgin olive oil",
        "Honeycrisp apples (4 count)",
    ]

    for item in grocery_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUTPUT)
    print(f"Initial file created: {OUTPUT}")

    # Open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print("GUI_READY: launched LibreOffice Writer with DISPLAY=:0")


create_initial()
