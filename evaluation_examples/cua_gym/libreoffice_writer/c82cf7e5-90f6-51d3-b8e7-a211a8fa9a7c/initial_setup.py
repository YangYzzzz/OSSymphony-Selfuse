"""
Initial Setup: Safety procedures document with DANGER, CAUTION, NOTICE keywords
Task ID: writer_txtfmt_062
Domain: libreoffice_writer

Creates a safety procedures document with 10 paragraphs.
DANGER appears in paragraph 2, CAUTION in paragraph 5, NOTICE in paragraph 8.
All text is in regular 12pt Arial black, no special formatting or character styles.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_txtfmt_062'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font for the document
    from docx.oxml.ns import qn
    from lxml import etree

    # Set document default font to Arial 12pt
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run('WORKPLACE SAFETY PROCEDURES MANUAL')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 1: Introduction
    p1 = doc.add_paragraph()
    run = p1.add_run(
        'This document outlines the essential safety procedures that all employees must follow '
        'when operating equipment and performing tasks within our facility. Adherence to these '
        'guidelines is mandatory and helps ensure a safe working environment for everyone.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 2: Contains DANGER
    p2 = doc.add_paragraph()
    run = p2.add_run(
        'DANGER: High voltage electrical panels are located throughout the facility. '
        'Only certified electricians are permitted to open or service these panels. '
        'Unauthorized access to electrical systems poses a severe risk of electrocution and death. '
        'All employees must maintain a minimum distance of three feet from any open electrical panel.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 3: Chemical Handling
    p3 = doc.add_paragraph()
    run = p3.add_run(
        'Chemical handling procedures must be followed at all times in the laboratory and storage areas. '
        'Personal protective equipment including gloves, safety goggles, and lab coats must be worn '
        'when handling any chemical substances. Material Safety Data Sheets (MSDS) for all chemicals '
        'are available in the safety office and online portal.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 4: Equipment Safety
    p4 = doc.add_paragraph()
    run = p4.add_run(
        'All machinery and equipment must be inspected before use. Employees should check for '
        'visible damage, loose parts, or unusual sounds that may indicate a malfunction. '
        'Equipment found to be in unsafe condition must be tagged out of service immediately '
        'using the lockout/tagout procedure and reported to the maintenance department.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 5: Contains CAUTION
    p5 = doc.add_paragraph()
    run = p5.add_run(
        'CAUTION: The floor surfaces in the loading dock and warehouse areas may become slippery '
        'when wet. All personnel entering these areas must wear slip-resistant footwear as per '
        'company policy. Wet floor signs must be placed immediately whenever spills occur, and '
        'spills should be cleaned up within five minutes to prevent accidents.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 6: Fire Safety
    p6 = doc.add_paragraph()
    run = p6.add_run(
        'Fire safety is paramount in all areas of our facility. Fire extinguishers are located '
        'at marked stations throughout each floor and must never be obstructed. Employees are '
        'required to complete annual fire safety training and must know the location of the '
        'nearest emergency exit from their workstation.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 7: Ergonomics
    p7 = doc.add_paragraph()
    run = p7.add_run(
        'Ergonomic guidelines have been established to prevent repetitive strain injuries and '
        'musculoskeletal disorders. Employees who work at computer workstations for extended '
        'periods should ensure their monitor is at eye level, chair is properly adjusted, '
        'and take regular breaks every 45 minutes to stretch and rest their eyes.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 8: Contains NOTICE
    p8 = doc.add_paragraph()
    run = p8.add_run(
        'NOTICE: All workplace accidents, injuries, and near-miss incidents must be reported '
        'to the Human Resources department within 24 hours of occurrence. Incident report forms '
        'are available at each department supervisor\'s desk and on the company intranet. '
        'Prompt reporting ensures timely medical attention and helps prevent future incidents.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 9: Emergency Procedures
    p9 = doc.add_paragraph()
    run = p9.add_run(
        'Emergency evacuation procedures are posted at each exit throughout the facility. '
        'In the event of an emergency alarm, all employees must immediately stop work, '
        'secure any sensitive materials, and proceed calmly to the nearest emergency exit. '
        'The designated assembly point is located in the north parking lot, section B.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 10: Compliance
    p10 = doc.add_paragraph()
    run = p10.add_run(
        'Compliance with these safety procedures is a condition of employment. Violations of '
        'safety protocols may result in disciplinary action up to and including termination. '
        'Employees are encouraged to report safety concerns to their supervisor or the Safety '
        'Committee without fear of retaliation. Safety is everyone\'s responsibility.'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Save the document
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Also copy to Desktop as the task references it there
    desktop_path = f'{WORKDIR}/Desktop/safety_procedures.docx'
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)
    import shutil
    shutil.copy(OUTPUT, desktop_path)
    print(f'Also copied to Desktop: {desktop_path}')

    # GUI-ready startup: open the initial file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
