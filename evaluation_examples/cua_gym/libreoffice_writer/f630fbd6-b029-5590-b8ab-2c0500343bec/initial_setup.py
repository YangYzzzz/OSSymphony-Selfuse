"""
Initial Setup: Product catalog document with image (6cm x 4cm), 'Parallel' text wrap, 0cm spacing
Task ID: writer_obj_011
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from PIL import Image as PILImage

WORKDIR = '/home/user'
TASK_ID = 'writer_obj_011'
OUTPUT = f'{WORKDIR}/Desktop/product_catalog.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_placeholder_image(width_px=240, height_px=160):
    """Create a simple product image in memory."""
    img = PILImage.new('RGB', (width_px, height_px), color=(70, 130, 180))
    # Add a simple border effect
    from PIL import ImageDraw
    draw = ImageDraw.Draw(img)
    draw.rectangle([2, 2, width_px - 3, height_px - 3], outline=(255, 255, 255), width=3)
    draw.text((width_px // 2 - 30, height_px // 2 - 10), "Product", fill=(255, 255, 255))
    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def add_floating_image_parallel_wrap(doc, image_stream, width_emu, height_emu,
                                      dist_t=0, dist_b=0, dist_l=0, dist_r=0,
                                      pos_x_emu=457200, pos_y_emu=457200):
    """
    Add a floating image with 'Parallel' (wrapSquare) text wrapping.
    dist_t/b/l/r: wrapping distance in EMU (0 = default, 0cm spacing).
    Inserts an <wp:anchor> element with <wp:wrapSquare>.
    """
    # Add the image part to the document
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    # We'll add image via a paragraph run, then replace the inline with an anchor
    para = doc.add_paragraph()

    # Add picture inline first to get the image relationship registered
    run = para.add_run()
    pic = run.add_picture(image_stream, width=width_emu, height=height_emu)

    # The inline drawing element is: run -> w:drawing -> wp:inline -> ...
    # We need to replace wp:inline with wp:anchor
    drawing_elem = run._r.find(qn('w:drawing'))
    inline_elem = drawing_elem.find(
        '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
    )

    if inline_elem is None:
        return  # fallback

    # Get the graphic element from inline
    extent = inline_elem.find(
        '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'
    )
    effect_extent = inline_elem.find(
        '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}effectExtent'
    )
    doc_pr = inline_elem.find(
        '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr'
    )
    graphic = inline_elem.find(
        '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}graphic'
    )
    if graphic is None:
        graphic = inline_elem.find(
            '{http://www.w3.org/XML/1998/namespace}graphic'
        )
    # The graphic is in the a: namespace
    graphic_elem = inline_elem.find(
        '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}graphic'
    )

    WPD = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

    # Build anchor XML string
    anchor_xml = f'''<wp:anchor
        distT="{dist_t}"
        distB="{dist_b}"
        distL="{dist_l}"
        distR="{dist_r}"
        simplePos="0"
        relativeHeight="251658240"
        behindDoc="0"
        locked="0"
        layoutInCell="1"
        allowOverlap="1"
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
        xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
            <wp:posOffset>{pos_x_emu}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>{pos_y_emu}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapSquare wrapText="bothSides" distT="{dist_t}" distB="{dist_b}" distL="{dist_l}" distR="{dist_r}"/>
        <wp:docPr id="1" name="Image 1" descr="Product Image"/>
        <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
    </wp:anchor>'''

    anchor_elem = etree.fromstring(anchor_xml)

    # Copy the graphic data from inline into anchor
    # The graphic is nested inside inline under wp:graphic (but actually it's in a: namespace)
    # Let's get all children of inline and reconstruct
    children_to_copy = list(inline_elem)

    # Find the a:graphic element (it's not under wp: namespace)
    for child in children_to_copy:
        tag = etree.QName(child.tag).localname
        if tag == 'graphic':
            # Copy this graphic element into the anchor
            anchor_elem.append(child)
            break

    # Replace inline with anchor in drawing element
    drawing_elem.remove(inline_elem)
    drawing_elem.append(anchor_elem)

    return para


def create_initial():
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # --- Page 1: Product Catalog Introduction ---
    title = doc.add_heading('TechVision Product Catalog 2025', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'Welcome to TechVision\'s comprehensive product catalog. '
        'Our range of innovative technology solutions is designed to meet '
        'the needs of modern businesses. From enterprise software to cutting-edge '
        'hardware, we offer products that drive efficiency and growth.'
    )

    doc.add_paragraph()

    # Section heading
    doc.add_heading('Featured Products', level=1)

    # Paragraph before image (to wrap around)
    p1 = doc.add_paragraph(
        'The ProSeries X500 is our flagship workstation, designed for professionals '
        'who demand peak performance. Equipped with the latest multi-core processors '
        'and high-speed NVMe storage, it handles the most demanding workloads with ease. '
        'The system supports up to 128GB DDR5 RAM and features dual Thunderbolt 4 ports '
        'for seamless peripheral connectivity.'
    )

    # Create product image (6cm x 4cm = 2160000 x 1440000 EMU) with 0cm wrap spacing
    width_emu = int(Cm(6))   # 2160000
    height_emu = int(Cm(4))  # 1440000

    image_stream = create_placeholder_image(240, 160)

    # Add floating image with Parallel (wrapSquare) wrap, 0cm spacing on all sides
    add_floating_image_parallel_wrap(
        doc, image_stream,
        width_emu=width_emu,
        height_emu=height_emu,
        dist_t=0, dist_b=0, dist_l=0, dist_r=0,
        pos_x_emu=457200,   # ~1.27cm from left
        pos_y_emu=914400,   # ~2.54cm from top
    )

    p2 = doc.add_paragraph(
        'The ProSeries X500 comes with a 3-year on-site warranty and dedicated '
        '24/7 technical support. It is certified for use in regulated industries '
        'including healthcare, finance, and government. The optional enterprise '
        'management suite allows IT administrators to remotely monitor and configure '
        'devices across the organization.'
    )

    p3 = doc.add_paragraph(
        'Pricing starts at $2,499 for the base configuration. Volume discounts are '
        'available for orders of 10 or more units. Contact our sales team for '
        'customized configurations and enterprise licensing options.'
    )

    # Add a page break to ensure content stays on page 1
    doc.add_page_break()

    # --- Page 2: Product Specifications ---
    doc.add_heading('Product Specifications', level=1)

    spec_intro = doc.add_paragraph(
        'Below are the detailed technical specifications for the ProSeries X500 workstation. '
        'All configurations include our standard 3-year warranty.'
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Specification'

    specs = [
        ('Processor', 'Intel Core i9-14900K, 24-core, 5.8GHz'),
        ('Memory', '64GB DDR5-6000 (4x16GB), expandable to 128GB'),
        ('Storage', '2TB NVMe SSD (PCIe 5.0) + 4TB HDD'),
        ('Graphics', 'NVIDIA RTX 4080 16GB GDDR6X'),
        ('Display Output', '4x DisplayPort 2.1, 2x HDMI 2.1'),
        ('Connectivity', '2x Thunderbolt 4, 8x USB 3.2, 2x USB-C'),
        ('Network', '10GbE LAN, Wi-Fi 6E, Bluetooth 5.3'),
        ('Power Supply', '1000W 80+ Platinum'),
        ('Dimensions', '46cm x 22cm x 48cm (H x W x D)'),
        ('Weight', '18.5 kg'),
        ('Operating System', 'Windows 11 Pro / Ubuntu 22.04 LTS'),
        ('Price', '$2,499 (base) - $4,299 (fully loaded)'),
    ]

    for spec_name, spec_value in specs:
        row_cells = table.add_row().cells
        row_cells[0].text = spec_name
        row_cells[1].text = spec_value

    doc.add_paragraph()
    doc.add_paragraph(
        'For additional product information, warranty details, and purchasing options, '
        'please visit our website at www.techvision.com or contact your regional '
        'sales representative.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
