"""
FINAL REWARD SCRIPT - SUCCESS
Task: In my LibreOffice Writer document, I want the text in paragraph 8 to switch columns part-way through. Right after the second sentence ends (that’s the period after “...quarter results.”), I need the rest of that paragraph to start in the next column. What’s the quickest way to drop in a manual column break exactly there?
Generated: 2025-09-10 15:25:42
Status: success
Model: azure-o3
Total Steps: 4
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn


def verify_manual_column_break(file_path: str) -> float:
    """Reward-function for the Writer task:

    Task to verify
    --------------
    1. A manual *column* break must exist **inside paragraph 8** (index 7).
    2. The break must appear **immediately after the second sentence**, i.e. after the period in
       “…first quarter results.”

    Scoring rules (progressive)
    ---------------------------
      • 0.6  – Column break of type "column" found in paragraph 8.
      • 0.4  – Break is positioned directly after the required sentence.
      • 1.0  – Both conditions satisfied.
    """

    print(f"Starting verification for file: {file_path}\n")

    # --- Prerequisites: file must load ------------------------------------------------
    if not os.path.exists(file_path):
        print("✗ File not found – task failed")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Could not open DOCX – {exc}")
        print("REWARD: 0.0")
        return 0.0

    # Ensure paragraph 8 exists --------------------------------------------------------
    target_idx = 7  # 0-based index for paragraph 8
    if len(doc.paragraphs) <= target_idx:
        print(f"✗ Document has only {len(doc.paragraphs)} paragraph(s); paragraph 8 missing")
        print("REWARD: 0.0")
        return 0.0

    para = doc.paragraphs[target_idx]

    # -------------------------------------------------------------------------
    # Requirement 1 ‑ detect a <w:br w:type="column"/> inside paragraph 8
    # -------------------------------------------------------------------------
    column_break_found = False
    for run in para.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "column":
                column_break_found = True
                break
        if column_break_found:
            break

    score = 0.0
    if column_break_found:
        print("✓ Column break detected in paragraph 8 (0.6 points)")
        score += 0.6
    else:
        print("✗ No column break found in paragraph 8 – task incomplete")
        print(f"REWARD: {score}")
        return score  # Cannot continue without the break

    # -------------------------------------------------------------------------
    # Requirement 2 ‑ verify correct position of the break
    # -------------------------------------------------------------------------
    before_text, after_text = "", ""
    encountered_break = False

    # Walk through paragraph XML collecting text before and after the break
    for child in para._p.iter():
        if child.tag == qn("w:br") and child.get(qn("w:type")) == "column":
            encountered_break = True
            continue
        if child.tag == qn("w:t"):
            if encountered_break:
                after_text += child.text or ""
            else:
                before_text += child.text or ""

    before_text = before_text.strip()
    after_text = after_text.strip()

    print("Text immediately before break:")
    print(before_text[:120] + ("..." if len(before_text) > 120 else ""))
    print("Text immediately after break:")
    print(after_text[:120] + ("..." if len(after_text) > 120 else ""))

    expected_ending = "quarter results."
    correct_location = re.search(re.escape(expected_ending) + r"\s*$", before_text, re.IGNORECASE) is not None

    if correct_location:
        print("✓ Column break positioned correctly after the second sentence (0.4 points)")
        score += 0.4
    else:
        print("✗ Column break exists but is not at the required location (0 points)")

    final_score = min(score, 1.0)
    print(f"REWARD: {final_score}")
    return final_score


# -----------------------------------------------------------------------------
# Execute when run as a script (so the autograder triggers the check)
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    DOC_PATH = "/home/user/in_my_libreoffice_writer_document_i_want_the_text_in_paragraph_8_to_switch_columns_part_way_through_.docx"
    verify_manual_column_break(DOC_PATH)

