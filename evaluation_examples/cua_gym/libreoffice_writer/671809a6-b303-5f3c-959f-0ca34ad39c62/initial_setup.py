"""
Initial Setup: CONFIDENTIAL document with normal character spacing on title
Task ID: writer_txtfmt_032
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKDIR = '/home/user/Desktop'  # Task context says file is on ~/Desktop
TASK_ID = 'writer_txtfmt_032'
OUTPUT = f'{WORKDIR}/classified_report.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title paragraph: CONFIDENTIAL ---
    # 16pt Arial Bold, normal character spacing (no condensed/expanded)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('CONFIDENTIAL')
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    # Normal character spacing = no w:spacing set in rPr (default = 0)

    # --- Four paragraphs of internal investigation findings ---
    para1 = doc.add_paragraph()
    run1 = para1.add_run(
        'Executive Summary: This report presents the findings of an internal investigation '
        'conducted between January 15 and February 28, 2025, regarding unauthorized data access '
        'in the Procurement Division. The investigation was initiated following anomalous '
        'system log entries identified by the IT Security team on January 12, 2025.'
    )
    run1.font.name = 'Arial'
    run1.font.size = Pt(11)

    para2 = doc.add_paragraph()
    run2 = para2.add_run(
        'Investigation Scope and Methodology: The investigation team, comprising three senior '
        'auditors and two external cybersecurity consultants, reviewed access logs spanning a '
        'six-month period from July 1 to December 31, 2024. A total of 847 access events were '
        'examined, with 23 flagged as potentially non-compliant with existing data governance '
        'policies. Interviews were conducted with 14 employees across five departments.'
    )
    run2.font.name = 'Arial'
    run2.font.size = Pt(11)

    para3 = doc.add_paragraph()
    run3 = para3.add_run(
        'Key Findings: The investigation determined that three individuals accessed vendor '
        'financial records outside their authorized scope. Specifically, records belonging to '
        'suppliers Meridian Logistics (Vendor ID: ML-4492) and Stratford Supply Co. '
        '(Vendor ID: SS-7731) were accessed without documented business justification on '
        'November 3, 8, and 19, 2024. No evidence of data exfiltration was found; however, '
        'the accesses represent a breach of the Least Privilege Access Policy (LPAP-2021).'
    )
    run3.font.name = 'Arial'
    run3.font.size = Pt(11)

    para4 = doc.add_paragraph()
    run4 = para4.add_run(
        'Recommendations and Corrective Actions: The investigation team recommends immediate '
        'review and tightening of role-based access controls within the procurement management '
        'system. Additionally, mandatory refresher training on data governance policies should '
        'be scheduled for all Procurement Division staff by March 31, 2025. The three '
        'individuals involved have been issued formal written warnings in accordance with '
        'HR Policy 4.7. A follow-up audit is recommended for Q3 2025 to verify compliance.'
    )
    run4.font.name = 'Arial'
    run4.font.size = Pt(11)

    # Ensure Desktop directory exists on VM (it should, but be safe)
    os.makedirs(WORKDIR, exist_ok=True)
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup: open the file in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
