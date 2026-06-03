"""
Reward Script: Alternating row colors on requirements table
Task ID: writer_tech_042
Domain: libreoffice_writer
Scoring:
  Component 1 (0.4): Light gray (#F5F5F5) shading exists on at least one row
  Component 2 (0.3): Rows alternate between two colors (white and light gray)
  Component 3 (0.3): All cells within each row have uniform shading
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_042'


def get_row_shadings(table):
    """Extract shading fill color for every cell in every row."""
    result = []
    for row in table.rows:
        row_fills = []
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            shd = tcPr.find(qn('w:shd')) if tcPr is not None else None
            fill = shd.get(qn('w:fill')) if shd is not None else None
            row_fills.append(fill)
        result.append(row_fills)
    return result


def normalize_color(c):
    """Normalize a color string to uppercase, treating None as absent."""
    if c is None:
        return None
    return c.upper().strip()


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]
    num_rows = len(table.rows)
    if num_rows < 2:
        print(f"FAIL: Table has only {num_rows} rows, need at least 2 for alternation")
        print("REWARD: 0.0")
        return 0.0

    shadings = get_row_shadings(table)

    # Component 1: Light gray (#F5F5F5) shading exists on at least one row (0.4 points)
    # This checks that F5F5F5 was applied — initial has all None, so this fails on initial.
    try:
        gray_cells = sum(
            1 for row_fills in shadings
            for fill in row_fills
            if normalize_color(fill) == 'F5F5F5'
        )

        if gray_cells > 0:
            print(f"PASS: Component 1 — Light gray #F5F5F5 shading found in table (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — No cell with #F5F5F5 shading found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Consecutive rows alternate between two colors (0.3 points)
    # Check that each pair of adjacent rows has different dominant shading colors,
    # using exactly two colors across all rows. Header row may match row 1.
    # Initial has all None — no alternation pattern, so this fails on initial.
    try:
        from collections import Counter
        # Get dominant color per row (most common non-None fill)
        row_colors = []
        for ri, row_fills in enumerate(shadings):
            normalized = [normalize_color(f) for f in row_fills]
            non_none = [c for c in normalized if c is not None]
            if non_none:
                dominant = Counter(non_none).most_common(1)[0][0]
                row_colors.append(dominant)
            else:
                row_colors.append(None)

        # Check alternation: need at least 2 distinct non-None colors
        distinct_colors = set(c for c in row_colors if c is not None)
        if len(distinct_colors) < 2:
            print(f"FAIL: Component 2 — Need at least 2 distinct row colors, found: {distinct_colors}")
        else:
            # Check that data rows (1 onwards) alternate; header (row 0) can be either color.
            # For data rows: consecutive rows must differ in color.
            fail_details = []
            for i in range(1, num_rows - 1):
                c_cur = row_colors[i]
                c_next = row_colors[i + 1]
                if c_cur is None or c_next is None:
                    fail_details.append(f"Row {i}={c_cur}, Row {i+1}={c_next} (None present)")
                elif c_cur == c_next:
                    fail_details.append(f"Row {i} and Row {i+1} both {c_cur}")

            if len(fail_details) == 0:
                print(f"PASS: Component 2 — Consecutive data rows alternate colors: {distinct_colors} (0.3 pts)")
                total_score += 0.3
            else:
                print(f"FAIL: Component 2 — Alternation breaks: {fail_details}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: All cells within each row have uniform shading (0.3 points)
    # Every cell in a row must have the same fill color (not None).
    # Initial has all None — so this check for "non-None and uniform" fails on initial.
    try:
        uniform_count = 0
        total_checked = 0
        for ri, row_fills in enumerate(shadings):
            normalized = [normalize_color(f) for f in row_fills]
            non_none = [c for c in normalized if c is not None]
            total_checked += 1
            # All cells must have shading and be the same color
            if len(non_none) == len(row_fills) and len(set(non_none)) == 1:
                uniform_count += 1
            else:
                print(f"  Row {ri}: not uniform — fills={normalized}")

        if total_checked > 0 and uniform_count == total_checked:
            print(f"PASS: Component 3 — All {uniform_count} rows have uniform cell shading (0.3 pts)")
            total_score += 0.3
        else:
            print(f"FAIL: Component 3 — {uniform_count}/{total_checked} rows have uniform shading")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
