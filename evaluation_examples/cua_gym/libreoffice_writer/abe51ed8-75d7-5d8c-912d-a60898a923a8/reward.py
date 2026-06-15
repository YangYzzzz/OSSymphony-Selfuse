"""
Reward Script: Replace bracketed references with endnotes
Task ID: writer_tech_083
Domain: libreoffice_writer
Scoring:
  - Component 1 (0.25): Bracket refs [1],[2],[3] removed from body text
  - Component 2 (0.15): Manual References section removed
  - Component 3 (0.25): Endnotes XML contains 3 user endnotes with correct content
  - Component 4 (0.20): 3 endnoteReference marks present in document body
  - Component 5 (0.15): Endnote content matches original reference texts
"""

import os
import re
import zipfile

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_083'


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Load document via python-docx for text checks
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load document {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Load raw XML for endnote structure checks
    try:
        zf = zipfile.ZipFile(file_path, 'r')
    except Exception as e:
        print(f"CRITICAL: Cannot open zip {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Bracket references [1],[2],[3] removed from body text (0.25 pts)
    # Only check body paragraphs (not the References section itself).
    # In the initial file, body paragraphs contain [1], [2], [3]. These must be gone.
    try:
        bracket_refs_found = []
        for i, para in enumerate(doc.paragraphs):
            # Skip any paragraph that is part of a references section
            if para.style.name.startswith('Heading') and 'reference' in para.text.lower():
                break
            matches = re.findall(r'\[\d+\]', para.text)
            if matches:
                bracket_refs_found.extend(matches)

        if len(bracket_refs_found) == 0:
            print(f"PASS: Component 1 — No bracket refs in body text (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 1 — Found bracket refs in body: {bracket_refs_found}")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Manual References section removed (0.15 pts)
    # Initial file has a "References" heading and 3 reference paragraphs at the end.
    # These must be removed in the golden file.
    try:
        ref_headings = [p for p in doc.paragraphs
                        if p.style.name.startswith('Heading') and p.text.strip().lower() == 'references']
        manual_ref_paras = [p for p in doc.paragraphs
                            if re.match(r'^\[\d+\]\s+\w+', p.text.strip())]
        has_references_heading = len(ref_headings) > 0
        has_manual_ref_para = len(manual_ref_paras) > 0

        if not has_references_heading and not has_manual_ref_para:
            print(f"PASS: Component 2 — No manual References section found (0.15 pts)")
            total_score += 0.15
        else:
            details = []
            if has_references_heading:
                details.append("References heading still present")
            if has_manual_ref_para:
                details.append("Manual ref paragraphs still present")
            print(f"FAIL: Component 2 — {'; '.join(details)}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Endnotes XML exists with 3 user endnotes (0.25 pts)
    # IDs 0 and 1 are separator/continuation; user endnotes have id >= 2
    try:
        if 'word/endnotes.xml' not in zf.namelist():
            print(f"FAIL: Component 3 — word/endnotes.xml does not exist")
        else:
            endnotes_xml = zf.read('word/endnotes.xml').decode('utf-8')
            # Count user endnotes (those without type="separator" or type="continuationSeparator")
            # User endnotes are <w:endnote w:id="N"> without w:type attribute
            all_endnotes = re.findall(r'<w:endnote\s+[^>]*w:id="(\d+)"[^>]*/?>',  endnotes_xml)
            separator_endnotes = re.findall(r'<w:endnote\s+[^>]*w:type="(?:separator|continuationSeparator)"[^>]*w:id="(\d+)"', endnotes_xml)
            # Also check reverse attribute order
            separator_endnotes2 = re.findall(r'<w:endnote\s+[^>]*w:id="(\d+)"[^>]*w:type="(?:separator|continuationSeparator)"', endnotes_xml)
            sep_ids = set(separator_endnotes) | set(separator_endnotes2)
            user_endnote_ids = [eid for eid in all_endnotes if eid not in sep_ids]

            if len(user_endnote_ids) >= 3:
                print(f"PASS: Component 3 — Found {len(user_endnote_ids)} user endnotes (ids: {user_endnote_ids}) (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 3 — Expected 3 user endnotes, found {len(user_endnote_ids)} (ids: {user_endnote_ids})")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: 3 endnoteReference marks in document body (0.20 pts)
    # These are <w:endnoteReference w:id="N"/> elements in word/document.xml
    try:
        doc_xml = zf.read('word/document.xml').decode('utf-8')
        endnote_ref_ids = re.findall(r'endnoteReference\s+w:id="(\d+)"', doc_xml)
        # Filter out separator references (ids 0, 1)
        user_ref_ids = [rid for rid in endnote_ref_ids if int(rid) >= 2]

        if len(user_ref_ids) >= 3:
            print(f"PASS: Component 4 — Found {len(user_ref_ids)} endnoteReference marks in body (ids: {user_ref_ids}) (0.20 pts)")
            total_score += 0.20
        elif len(user_ref_ids) > 0:
            partial = round(0.20 * len(user_ref_ids) / 3, 2)
            print(f"PARTIAL: Component 4 — Found {len(user_ref_ids)}/3 endnoteReference marks (ids: {user_ref_ids}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 — No endnoteReference marks found in document body")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Endnote content matches reference texts (0.15 pts)
    # Check that the 3 key reference authors appear in endnotes
    try:
        if 'word/endnotes.xml' not in zf.namelist():
            print(f"FAIL: Component 5 — No endnotes.xml")
        else:
            endnotes_xml = zf.read('word/endnotes.xml').decode('utf-8')
            expected_authors = ['Ruder', 'Kingma', 'Loshchilov']
            found_count = 0
            for author in expected_authors:
                if author in endnotes_xml:
                    found_count += 1

            if found_count == 3:
                print(f"PASS: Component 5 — All 3 reference texts found in endnotes (0.15 pts)")
                total_score += 0.15
            elif found_count > 0:
                partial = round(0.15 * found_count / 3, 2)
                print(f"PARTIAL: Component 5 — Found {found_count}/3 reference authors in endnotes ({partial} pts)")
                total_score += partial
            else:
                print(f"FAIL: Component 5 — No reference texts found in endnotes")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    zf.close()

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
