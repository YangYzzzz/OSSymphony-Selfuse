"""
Initial Setup: Replace manually numbered references with proper endnotes
Task ID: writer_tech_083
Domain: libreoffice_writer

Creates a technical document with inline bracketed references [1], [2], [3]
and a manually typed reference list at the end. No endnotes present.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_tech_083'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Advances in Neural Network Optimization Techniques', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author info ---
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = author_para.add_run('Dr. Elena Rodriguez, Department of Computer Science')
    run.font.size = Pt(11)
    run.font.italic = True

    # --- Abstract ---
    doc.add_heading('Abstract', level=2)
    abstract = doc.add_paragraph(
        'This paper surveys recent developments in optimization algorithms for deep neural networks. '
        'We examine adaptive learning rate methods, regularization strategies, and second-order '
        'optimization approaches that have shown significant improvements in training convergence '
        'and generalization performance across a variety of benchmark tasks.'
    )
    abstract.paragraph_format.space_after = Pt(12)

    # --- Introduction with inline references ---
    doc.add_heading('1. Introduction', level=2)
    intro1 = doc.add_paragraph(
        'Deep learning has transformed numerous fields including computer vision, natural language '
        'processing, and speech recognition. The success of these models depends critically on '
        'the optimization algorithms used during training. Stochastic gradient descent (SGD) and '
        'its variants remain the workhorses of neural network training [1], but significant progress '
        'has been made in developing more sophisticated methods.'
    )
    intro1.paragraph_format.space_after = Pt(6)

    intro2 = doc.add_paragraph(
        'Adaptive learning rate methods, such as Adam and its successors, have gained widespread '
        'adoption due to their ability to automatically adjust per-parameter learning rates [2]. '
        'These methods combine ideas from momentum-based optimization with second-moment estimation '
        'to achieve faster convergence on many practical problems.'
    )
    intro2.paragraph_format.space_after = Pt(6)

    # --- Methods section with another reference ---
    doc.add_heading('2. Optimization Methods', level=2)
    methods1 = doc.add_paragraph(
        'Recent work has focused on bridging the generalization gap between adaptive methods and '
        'well-tuned SGD with momentum. Techniques such as learning rate warmup, gradient clipping, '
        'and decoupled weight decay have been shown to improve the performance of adaptive optimizers '
        'significantly [3]. In particular, the AdamW variant, which decouples the weight decay from '
        'the gradient update, has become the default optimizer for transformer-based architectures.'
    )
    methods1.paragraph_format.space_after = Pt(6)

    methods2 = doc.add_paragraph(
        'Second-order methods, while computationally expensive, offer the advantage of curvature '
        'information that can lead to better optimization trajectories. Approximations such as '
        'K-FAC and Shampoo have made these approaches practical for large-scale models by reducing '
        'the per-step computational overhead to near first-order levels.'
    )
    methods2.paragraph_format.space_after = Pt(6)

    # --- Results section ---
    doc.add_heading('3. Experimental Results', level=2)
    results = doc.add_paragraph(
        'Our experiments across CIFAR-100, ImageNet, and WMT-14 translation benchmarks demonstrate '
        'that the combination of adaptive learning rates with modern regularization techniques '
        'achieves state-of-the-art performance. The AdamW optimizer with cosine annealing '
        'consistently outperformed vanilla SGD by 1.2-2.8% in top-1 accuracy on image '
        'classification tasks, while matching or exceeding SGD generalization on language tasks.'
    )
    results.paragraph_format.space_after = Pt(6)

    # --- Conclusion ---
    doc.add_heading('4. Conclusion', level=2)
    conclusion = doc.add_paragraph(
        'The optimization landscape for deep learning continues to evolve rapidly. While no single '
        'method dominates across all settings, adaptive methods with proper regularization represent '
        'the current best practice for most practical applications. Future work should explore the '
        'theoretical foundations of these empirical observations and develop unified frameworks '
        'that combine the strengths of different optimization paradigms.'
    )
    conclusion.paragraph_format.space_after = Pt(12)

    # --- Manual Reference List ---
    doc.add_heading('References', level=2)
    ref1 = doc.add_paragraph(
        '[1] Ruder, S. (2016). An overview of gradient descent optimization algorithms. '
        'arXiv preprint arXiv:1609.04747.'
    )
    ref1.paragraph_format.space_after = Pt(4)

    ref2 = doc.add_paragraph(
        '[2] Kingma, D.P. and Ba, J. (2015). Adam: A method for stochastic optimization. '
        'Proceedings of the 3rd International Conference on Learning Representations (ICLR).'
    )
    ref2.paragraph_format.space_after = Pt(4)

    ref3 = doc.add_paragraph(
        '[3] Loshchilov, I. and Hutter, F. (2019). Decoupled weight decay regularization. '
        'Proceedings of the 7th International Conference on Learning Representations (ICLR).'
    )
    ref3.paragraph_format.space_after = Pt(4)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
