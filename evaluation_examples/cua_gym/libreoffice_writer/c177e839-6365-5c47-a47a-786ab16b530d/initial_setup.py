"""
Initial Setup: Employee handbook document with section headings but no bookmarks.
Task ID: writer_hr_055
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_055'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # -- Title --
    title = doc.add_heading('Horizon Technologies Employee Handbook', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Effective Date: January 1, 2025')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ===== Section 1: Introduction =====
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Welcome to Horizon Technologies! We are thrilled to have you as part of our growing team. '
        'This employee handbook is designed to provide you with important information about our company '
        'policies, procedures, and the benefits available to you as an employee.'
    )
    doc.add_paragraph(
        'Horizon Technologies was founded in 2012 with a mission to deliver innovative software solutions '
        'to enterprises worldwide. Today, we employ over 1,200 professionals across offices in San Francisco, '
        'Austin, London, and Singapore. Our commitment to excellence drives everything we do.'
    )
    doc.add_paragraph(
        'This handbook applies to all full-time, part-time, and contract employees. It is not intended '
        'to create a contract of employment and does not alter the at-will nature of your employment. '
        'We reserve the right to modify these policies at any time with reasonable notice.'
    )

    # ===== Section 2: Employment Policies =====
    doc.add_heading('Employment Policies', level=1)
    doc.add_paragraph(
        'Horizon Technologies is an equal opportunity employer. We are committed to providing a work '
        'environment free from discrimination based on race, color, religion, sex, national origin, age, '
        'disability, veteran status, sexual orientation, gender identity, or any other protected characteristic.'
    )
    doc.add_paragraph(
        'All employment decisions, including hiring, promotions, transfers, and terminations, are based on '
        'legitimate business needs, job requirements, and individual qualifications. Employees are expected '
        'to maintain professional conduct at all times and adhere to the standards outlined in this handbook.'
    )
    doc.add_paragraph(
        'New employees undergo a 90-day probationary period during which performance and cultural fit '
        'are evaluated. During this period, either the employee or the company may terminate the employment '
        'relationship with one week\'s written notice. Regular performance reviews are conducted semi-annually '
        'in June and December.'
    )

    # ===== Section 3: Compensation =====
    doc.add_heading('Compensation', level=1)
    doc.add_paragraph(
        'Horizon Technologies offers competitive compensation packages benchmarked against industry standards. '
        'Salaries are reviewed annually in March, and adjustments are based on individual performance, market '
        'conditions, and company financial performance.'
    )
    doc.add_paragraph(
        'Employees are paid on a bi-weekly basis via direct deposit. Pay periods run from Monday through the '
        'following Sunday, with paychecks issued the Friday after each pay period closes. Overtime for '
        'non-exempt employees is compensated at 1.5 times the regular hourly rate for hours worked beyond 40 '
        'in a workweek.'
    )
    doc.add_paragraph(
        'Annual performance bonuses are awarded at the discretion of management and typically range from '
        '5% to 20% of base salary. The company also maintains a stock option plan for employees at the '
        'senior associate level and above, with a four-year vesting schedule and a one-year cliff.'
    )

    # ===== Section 4: Benefits =====
    doc.add_heading('Benefits', level=1)
    doc.add_paragraph(
        'Full-time employees are eligible for a comprehensive benefits package beginning on the first day of '
        'the month following 30 days of continuous employment. Our benefits include medical, dental, and '
        'vision insurance through Aetna and Delta Dental.'
    )
    doc.add_paragraph(
        'The company contributes 80% of the premium for individual coverage and 60% for dependent coverage. '
        'Employees may also participate in our 401(k) retirement plan with a company match of up to 4% of '
        'annual salary. Additional benefits include life insurance at two times annual salary, short-term '
        'and long-term disability coverage, and an Employee Assistance Program (EAP).'
    )
    doc.add_paragraph(
        'We also offer a $2,500 annual professional development stipend, a $75 monthly wellness reimbursement, '
        'and tuition reimbursement of up to $5,250 per calendar year for approved courses and certifications. '
        'Commuter benefits and flexible spending accounts (FSA) are available through pre-tax payroll deductions.'
    )

    # ===== Section 5: Leave Policies =====
    doc.add_heading('Leave Policies', level=1)
    doc.add_paragraph(
        'Horizon Technologies provides generous leave policies to support work-life balance. Full-time employees '
        'accrue 15 days of paid time off (PTO) per year during their first three years of employment, increasing '
        'to 20 days after the third anniversary and 25 days after the seventh anniversary.'
    )
    doc.add_paragraph(
        'In addition to PTO, the company observes 10 paid holidays per year: New Year\'s Day, Martin Luther '
        'King Jr. Day, Presidents\' Day, Memorial Day, Independence Day, Labor Day, Columbus Day, Veterans Day, '
        'Thanksgiving Day, and Christmas Day. Employees also receive two floating holidays per year.'
    )
    doc.add_paragraph(
        'Parental leave provides 12 weeks of fully paid leave for the birth, adoption, or foster placement of a '
        'child. Medical leave follows FMLA guidelines, providing up to 12 weeks of job-protected leave. '
        'Bereavement leave of up to 5 days is available for the loss of an immediate family member. Jury duty '
        'leave is provided with full pay for the duration of service.'
    )

    # ===== Section 6: Code of Conduct =====
    doc.add_heading('Code of Conduct', level=1)
    doc.add_paragraph(
        'All employees are expected to conduct themselves with integrity, professionalism, and respect for others. '
        'This Code of Conduct establishes the standards that guide our behavior and interactions in the workplace.'
    )
    doc.add_paragraph(
        'Employees must avoid conflicts of interest and disclose any potential conflicts to their manager or '
        'the Human Resources department. Confidential information, including trade secrets, client data, and '
        'proprietary technology, must be protected at all times. Unauthorized disclosure of confidential '
        'information is grounds for immediate termination.'
    )
    doc.add_paragraph(
        'Harassment, bullying, and retaliation of any kind are strictly prohibited. Employees who witness or '
        'experience such behavior should report it immediately through the company\'s Ethics Hotline '
        '(1-800-555-0199) or directly to Human Resources. All reports are treated confidentially, and no '
        'employee will face retaliation for making a good-faith report.'
    )

    # ===== Section 7: Safety =====
    doc.add_heading('Safety', level=1)
    doc.add_paragraph(
        'Horizon Technologies is committed to maintaining a safe and healthy work environment for all employees. '
        'We comply with all applicable federal, state, and local occupational safety and health regulations, '
        'including OSHA standards.'
    )
    doc.add_paragraph(
        'All employees are required to complete safety orientation training within their first week of employment. '
        'Department-specific safety training is conducted quarterly. Employees must report any unsafe conditions, '
        'equipment malfunctions, or workplace injuries to their supervisor and the Safety Committee within 24 hours.'
    )
    doc.add_paragraph(
        'Emergency evacuation procedures are posted in all common areas. Fire drills are conducted twice per year. '
        'First aid kits are located on each floor, and AED devices are available in the main lobby and cafeteria. '
        'The company provides ergonomic workstation assessments upon request to prevent repetitive strain injuries.'
    )

    # ===== Section 8: Acknowledgment =====
    doc.add_heading('Acknowledgment', level=1)
    doc.add_paragraph(
        'By signing below, I acknowledge that I have received and read the Horizon Technologies Employee '
        'Handbook. I understand that the policies described herein are guidelines and that the company reserves '
        'the right to change, modify, or discontinue any policy at any time.'
    )
    doc.add_paragraph(
        'I understand that this handbook does not constitute a contract of employment and that my employment '
        'with Horizon Technologies is at-will. Either the company or I may terminate the employment relationship '
        'at any time, with or without cause, and with or without notice.'
    )

    # Signature block
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run('Employee Name: ').bold = True
    sig.add_run('_' * 40)

    sig2 = doc.add_paragraph()
    sig2.add_run('Signature: ').bold = True
    sig2.add_run('_' * 40)

    sig3 = doc.add_paragraph()
    sig3.add_run('Date: ').bold = True
    sig3.add_run('_' * 40)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
