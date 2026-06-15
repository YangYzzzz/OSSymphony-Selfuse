"""
Reward Script: Enable 'Repeat heading rows' for inventory table
Task ID: writer_tm_016
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): tblHeader XML element present on first row
  Component 2 (0.3): Header row content is intact (SKU, Name, Category, Stock, Reorder Level, Supplier)
                      AND tblHeader is set (compound check anchored to the change)
  Component 3 (0.2): Table structure preserved (51 rows, 6 cols) AND tblHeader is set
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_016'

EXPECTED_HEADERS = ['SKU', 'Name', 'Category', 'Stock', 'Reorder Level', 'Supplier']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one table
    if len(doc.tables) == 0:
        print("FAIL: No tables found in the document")
        print("REWARD: 0.0")
        return 0.0

    table = doc.tables[0]

    # Check if tblHeader is set on row 0 (the core task change)
    row0_tr = table.rows[0]._tr
    trPr = row0_tr.find(qn('w:trPr'))
    has_tbl_header = False
    if trPr is not None:
        tbl_header_elems = trPr.findall(qn('w:tblHeader'))
        if len(tbl_header_elems) > 0:
            # Also check that val is not explicitly "false" / "0"
            elem = tbl_header_elems[0]
            val_attr = elem.get(qn('w:val'))
            # If val is absent, or "true", or "1", it means enabled
            if val_attr is None or val_attr.lower() in ('true', '1', 'on'):
                has_tbl_header = True

    # Component 1: tblHeader element present on first row (0.5 points)
    # This is the primary task change — fails on initial, passes on golden
    try:
        if has_tbl_header:
            print(f"PASS: Component 1 — tblHeader is set on first table row (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — tblHeader not found or disabled on first table row")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header row content intact AND tblHeader set (0.3 points)
    # Compound check: data integrity anchored to the task change
    try:
        actual_headers = [cell.text.strip() for cell in table.rows[0].cells]
        headers_match = actual_headers == EXPECTED_HEADERS
        if has_tbl_header and headers_match:
            print(f"PASS: Component 2 — Headers intact {actual_headers} AND tblHeader set (0.3 pts)")
            total_score += 0.3
        elif not has_tbl_header:
            print(f"FAIL: Component 2 — tblHeader not set (headers: {actual_headers})")
        else:
            print(f"FAIL: Component 2 — Headers mismatch: expected {EXPECTED_HEADERS}, got {actual_headers}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Table structure preserved (51 rows, 6 cols) AND tblHeader set (0.2 points)
    # Compound check: structural integrity anchored to the task change
    try:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        structure_ok = (num_rows == 51 and num_cols == 6)
        if has_tbl_header and structure_ok:
            print(f"PASS: Component 3 — Table structure {num_rows}x{num_cols} intact AND tblHeader set (0.2 pts)")
            total_score += 0.2
        elif not has_tbl_header:
            print(f"FAIL: Component 3 — tblHeader not set (table: {num_rows}x{num_cols})")
        else:
            print(f"FAIL: Component 3 — Table structure changed: expected 51x6, got {num_rows}x{num_cols}")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
