"""
Initial Setup: Enable 'Repeat heading rows' for inventory table
Task ID: writer_tm_016
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_016'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set up page margins for letter-size
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_heading('Warehouse Inventory List', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph(
        'This document contains the current inventory for the West Coast '
        'Distribution Center as of March 2025. All stock levels have been '
        'verified during the most recent physical count.'
    )
    intro.paragraph_format.space_after = Pt(12)

    # Create a 6-column table with 51 rows (1 header + 50 data)
    num_data_rows = 50
    table = doc.add_table(rows=1 + num_data_rows, cols=6)
    table.style = 'Table Grid'

    # Headers
    headers = ['SKU', 'Name', 'Category', 'Stock', 'Reorder Level', 'Supplier']
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # Realistic inventory data
    inventory_data = [
        ('WH-1001', 'Industrial Ball Bearings (10mm)', 'Hardware', 2450, 500, 'Pacific Fasteners Inc.'),
        ('WH-1002', 'Stainless Steel Bolts M8x30', 'Hardware', 8720, 2000, 'Midwest Metal Supply'),
        ('WH-1003', 'LED Panel Light 60W', 'Electrical', 340, 100, 'BrightTech Solutions'),
        ('WH-1004', 'Copper Wire 14 AWG (100ft)', 'Electrical', 1580, 300, 'Atlas Wire & Cable'),
        ('WH-1005', 'Hydraulic Hose 1/2" x 6ft', 'Plumbing', 620, 150, 'FlowMaster Hydraulics'),
        ('WH-1006', 'PVC Pipe Elbow 90deg 2"', 'Plumbing', 3200, 800, 'National Pipe Co.'),
        ('WH-1007', 'Safety Goggles Anti-Fog', 'Safety', 890, 200, 'Guardian Safety Gear'),
        ('WH-1008', 'Nitrile Gloves Large (Box/100)', 'Safety', 1450, 400, 'MedPro Supplies'),
        ('WH-1009', 'Cordless Drill 18V Kit', 'Power Tools', 175, 50, 'DeWalt Industrial'),
        ('WH-1010', 'Circular Saw Blade 7-1/4"', 'Power Tools', 540, 120, 'Freud America Inc.'),
        ('WH-1011', 'Welding Rod E6013 (5kg)', 'Welding', 320, 80, 'Lincoln Electric'),
        ('WH-1012', 'Argon Gas Cylinder 80cf', 'Welding', 45, 15, 'Airgas Specialty'),
        ('WH-1013', 'Concrete Anchor Bolt 3/8x3"', 'Construction', 6500, 1500, 'Hilti Corporation'),
        ('WH-1014', 'Drywall Screw #6x1-5/8"', 'Construction', 24000, 5000, 'Grip-Rite Fasteners'),
        ('WH-1015', 'Epoxy Adhesive 2-Part (gal)', 'Adhesives', 210, 60, 'Henkel Loctite'),
        ('WH-1016', 'Thread Locker Blue 50ml', 'Adhesives', 780, 200, 'Henkel Loctite'),
        ('WH-1017', 'Air Filter 20x25x1 MERV-8', 'HVAC', 1100, 300, 'FilterBuy Direct'),
        ('WH-1018', 'Refrigerant R-410A (25lb)', 'HVAC', 85, 25, 'Chemours Company'),
        ('WH-1019', 'Cable Tie 12" Black (1000pk)', 'Electrical', 920, 250, 'Panduit Corp.'),
        ('WH-1020', 'Junction Box 4x4 Metal', 'Electrical', 2300, 500, 'Hubbell Incorporated'),
        ('WH-1021', 'Sandpaper 220-Grit (50pk)', 'Abrasives', 640, 150, '3M Industrial'),
        ('WH-1022', 'Grinding Disc 4.5" Type 27', 'Abrasives', 1850, 400, 'Norton Abrasives'),
        ('WH-1023', 'Pipe Wrench 18" Heavy Duty', 'Hand Tools', 120, 30, 'Ridgid Tools'),
        ('WH-1024', 'Torque Wrench 1/2" Drive', 'Hand Tools', 95, 25, 'Snap-On Industrial'),
        ('WH-1025', 'Spray Paint Gloss Black (12oz)', 'Paint', 1650, 400, 'Rust-Oleum Corp.'),
        ('WH-1026', 'Primer White Interior (gal)', 'Paint', 430, 100, 'Sherwin-Williams'),
        ('WH-1027', 'Silicone Sealant Clear (10oz)', 'Sealants', 1920, 500, 'DAP Products'),
        ('WH-1028', 'Foam Insulation Board 4x8x1"', 'Insulation', 280, 75, 'Owens Corning'),
        ('WH-1029', 'Fiberglass Batt R-19 (16pk)', 'Insulation', 350, 90, 'CertainTeed Corp.'),
        ('WH-1030', 'Steel Shelf Unit 48x24x72"', 'Storage', 65, 20, 'Uline Inc.'),
        ('WH-1031', 'Pallet Jack 5500lb Cap.', 'Material Handling', 12, 5, 'Crown Equipment'),
        ('WH-1032', 'Stretch Wrap 18" (4 rolls)', 'Packaging', 520, 130, 'Uline Inc.'),
        ('WH-1033', 'Corrugated Box 24x18x12"', 'Packaging', 3800, 1000, 'Pratt Industries'),
        ('WH-1034', 'Fire Extinguisher ABC 10lb', 'Safety', 48, 15, 'Kidde Safety'),
        ('WH-1035', 'First Aid Kit 50-Person', 'Safety', 32, 10, 'Medline Industries'),
        ('WH-1036', 'Forklift Battery 36V 750AH', 'Material Handling', 8, 3, 'EnerSys Inc.'),
        ('WH-1037', 'Conveyor Belt 24" x 50ft', 'Material Handling', 15, 5, 'Dorner Mfg.'),
        ('WH-1038', 'Lubricant WD-40 (gal)', 'Maintenance', 290, 75, 'WD-40 Company'),
        ('WH-1039', 'Bearing Grease NLGI-2 (14oz)', 'Maintenance', 680, 150, 'Shell Lubricants'),
        ('WH-1040', 'Wire Rope 3/8" Galvanized', 'Rigging', 1100, 250, 'Loos & Co. Inc.'),
        ('WH-1041', 'Shackle 1/2" Screw Pin', 'Rigging', 450, 100, 'Crosby Group'),
        ('WH-1042', 'Pressure Gauge 0-300 PSI', 'Instrumentation', 220, 50, 'Ashcroft Inc.'),
        ('WH-1043', 'Thermocouple Type K 6"', 'Instrumentation', 380, 80, 'Omega Engineering'),
        ('WH-1044', 'Motor 5HP 3-Phase 1750RPM', 'Motors', 28, 10, 'Baldor Electric'),
        ('WH-1045', 'V-Belt A68 Industrial', 'Motors', 560, 120, 'Gates Corporation'),
        ('WH-1046', 'Pneumatic Cylinder 2" Bore', 'Pneumatics', 140, 35, 'Parker Hannifin'),
        ('WH-1047', 'Air Regulator 1/4" NPT', 'Pneumatics', 310, 70, 'SMC Corporation'),
        ('WH-1048', 'Stainless Tubing 1/2" OD 10ft', 'Tubing', 480, 100, 'Swagelok Company'),
        ('WH-1049', 'Rubber Gasket Set Assorted', 'Sealing', 720, 180, 'Garlock Sealing'),
        ('WH-1050', 'Chain Hoist 2-Ton Manual', 'Lifting', 18, 5, 'Harrington Hoists'),
    ]

    for i, (sku, name, cat, stock, reorder, supplier) in enumerate(inventory_data):
        row_idx = i + 1
        table.cell(row_idx, 0).text = sku
        table.cell(row_idx, 1).text = name
        table.cell(row_idx, 2).text = cat
        table.cell(row_idx, 3).text = str(stock)
        table.cell(row_idx, 4).text = str(reorder)
        table.cell(row_idx, 5).text = supplier

    # Explicitly ensure tblHeader is NOT set on the first row
    first_row = table.rows[0]
    trPr = first_row._tr.find(qn('w:trPr'))
    if trPr is not None:
        for hdr in trPr.findall(qn('w:tblHeader')):
            trPr.remove(hdr)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
