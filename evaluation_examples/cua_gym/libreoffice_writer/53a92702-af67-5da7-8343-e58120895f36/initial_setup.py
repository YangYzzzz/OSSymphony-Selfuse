"""
Initial Setup: Create a Writer document with four footnotes using default footnote separator.
Task ID: writer_bs_022
Domain: libreoffice_writer

Strategy:
  1. Create base .docx with python-docx (footnotes via XML)
  2. Open in LibreOffice via UNO
  3. Save as .odt using storeToURL with "writer8" filter (preserves ODF properties)
  4. Remove leftover .docx
  5. Kill LO, then relaunch with the .odt for GUI-ready state
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_022'
DOCX_PATH = f'{WORKDIR}/{TASK_ID}.docx'
ODT_PATH = f'{WORKDIR}/{TASK_ID}.odt'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def _add_footnote(doc, paragraph, number, text):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = paragraph.add_run()
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)
    run._element.insert(0, rPr)

    footnoteRef = OxmlElement('w:footnoteReference')
    footnoteRef.set(qn('w:id'), number)
    run._element.append(footnoteRef)

    _ensure_footnote(doc, number, text)


def _ensure_footnote(doc, fn_id, text):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree

    footnotes_part = None
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        footnotes_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )
        footnotes_element = etree.fromstring(footnotes_xml.encode('utf-8'))
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
        part_name = PackURI('/word/footnotes.xml')
        footnotes_part = Part(part_name, content_type, footnotes_xml.encode('utf-8'), doc.part.package)
        doc.part.relate_to(footnotes_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
        footnotes_part._element = footnotes_element
    else:
        footnotes_element = etree.fromstring(footnotes_part.blob)

    footnote = OxmlElement('w:footnote')
    footnote.set(qn('w:id'), fn_id)

    p = OxmlElement('w:p')

    r1 = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    r1.append(rPr)
    fRef = OxmlElement('w:footnoteRef')
    r1.append(fRef)
    p.append(r1)

    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.set(qn('xml:space'), 'preserve')
    t2.text = ' ' + text
    r2.append(t2)
    p.append(r2)

    footnote.append(p)
    footnotes_element.append(footnote)

    footnotes_part._blob = etree.tostring(footnotes_element, xml_declaration=True, encoding='UTF-8', standalone=True)


def create_base_document():
    """Create the base .docx document with four footnotes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    title = doc.add_heading('Quarterly Financial Performance Report', level=1)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('Prepared by the Finance Department \u2014 Q1 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    doc.add_heading('1. Executive Summary', level=2)

    p1 = doc.add_paragraph(
        'The first quarter of 2025 demonstrated resilient growth across all major business '
        'segments. Total revenue reached $128.4 million, representing a 12.3% increase '
        'year-over-year. Operating expenses were managed effectively, resulting in an '
        'operating margin of 23.7%.'
    )
    _add_footnote(doc, p1, '1', 'Revenue figures exclude intercompany transfers and are reported on an accrual basis per IFRS 15 guidelines.')

    doc.add_paragraph(
        'The technology services division led growth with a 18.2% revenue increase, '
        'driven primarily by enterprise cloud adoption and managed security services. '
        'Client retention rates improved to 94.6%, up from 91.2% in the previous quarter.'
    )

    doc.add_heading('2. Revenue Breakdown by Division', level=2)

    p3 = doc.add_paragraph(
        'Enterprise Solutions contributed $52.8 million (41.1% of total revenue), '
        'while Digital Transformation Services generated $38.6 million (30.1%). '
        'The remaining $37.0 million came from Consulting and Advisory services, '
        'which saw particularly strong demand in regulatory compliance engagements.'
    )
    _add_footnote(doc, p3, '2', 'Division revenue allocations follow the updated cost-center methodology adopted in January 2025, which reallocated shared infrastructure costs proportionally.')

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Division', 'Q1 2025 ($M)', 'Q4 2024 ($M)', 'Change (%)']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Enterprise Solutions', '$52.8', '$49.1', '+7.5%'],
        ['Digital Transformation', '$38.6', '$32.4', '+19.1%'],
        ['Consulting & Advisory', '$37.0', '$33.0', '+12.1%'],
        ['Total', '$128.4', '$114.5', '+12.3%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    doc.add_paragraph()

    doc.add_heading('3. Operating Expenses', level=2)

    p4 = doc.add_paragraph(
        'Total operating expenses for Q1 2025 were $97.9 million, a 9.8% increase '
        'from Q4 2024. Personnel costs accounted for 62.3% of operating expenses, '
        'reflecting the strategic hiring initiative launched in November 2024. '
        'Technology infrastructure costs rose by 14.1% due to expanded cloud capacity '
        'and data center consolidation efforts.'
    )
    _add_footnote(doc, p4, '3', 'Personnel costs include base salary, variable compensation, benefits, and equity-based compensation expense recognized under ASC 718.')

    doc.add_paragraph(
        'General and administrative expenses decreased by 3.2% following the '
        'implementation of automated procurement workflows and renegotiated vendor '
        'contracts. Travel and entertainment expenses were reduced by 11.7% through '
        'the expanded use of virtual meeting platforms across all regional offices.'
    )

    doc.add_heading('4. Forward Outlook', level=2)

    p6 = doc.add_paragraph(
        'Management expects Q2 2025 revenue to be in the range of $135\u2013$142 million, '
        'driven by the anticipated onboarding of three large enterprise contracts '
        'signed in late Q1. Capital expenditure for the full year is projected at '
        '$18.5 million, primarily allocated to the new East Coast data center and '
        'AI infrastructure investments.'
    )
    _add_footnote(doc, p6, '4', 'Forward-looking projections are based on current pipeline visibility and assume no material changes to macroeconomic conditions or regulatory landscape.')

    doc.add_paragraph(
        'The Board of Directors has approved an interim dividend of $0.42 per share, '
        'payable on June 15, 2025, to shareholders of record as of May 30, 2025. '
        'This represents a 5.0% increase over the prior interim dividend.'
    )

    doc.save(DOCX_PATH)
    print(f'Base .docx created: {DOCX_PATH}')


def convert_to_odt():
    """Open .docx in LibreOffice via UNO, save as .odt (writer8), then kill LO."""
    os.environ['DISPLAY'] = ':0'

    # Kill any running LO
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(3)

    # Start LO with UNO accept socket
    subprocess.Popen(
        ['soffice', '--writer', '--norestore',
         '--accept=socket,host=localhost,port=2002;urp;StarOffice.ComponentContext',
         DOCX_PATH],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(8)

    import uno
    from com.sun.star.beans import PropertyValue

    localContext = uno.getComponentContext()
    resolver = localContext.ServiceManager.createInstanceWithContext(
        'com.sun.star.bridge.UnoUrlResolver', localContext)

    ctx = None
    for attempt in range(20):
        try:
            ctx = resolver.resolve(
                'uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext')
            break
        except Exception as e:
            if attempt >= 19:
                raise RuntimeError(f'Could not connect to LO after 20 attempts: {e}')
            time.sleep(2)

    smgr = ctx.ServiceManager
    desktop = smgr.createInstanceWithContext('com.sun.star.frame.Desktop', ctx)

    # Wait for document to load
    doc = None
    for attempt in range(15):
        doc = desktop.getCurrentComponent()
        if doc is not None and hasattr(doc, 'StyleFamilies'):
            break
        time.sleep(2)

    if doc is None or not hasattr(doc, 'StyleFamilies'):
        raise RuntimeError('Document not loaded properly')

    print('Document loaded via UNO -- saving as .odt with default footnote properties')

    # Save as .odt using writer8 filter
    odt_url = 'file://' + ODT_PATH
    filter_prop = PropertyValue('FilterName', 0, 'writer8', 0)
    doc.storeToURL(odt_url, (filter_prop,))
    print(f'Saved as .odt: {ODT_PATH}')
    time.sleep(1)

    # Kill LO
    subprocess.run(['pkill', '-f', 'soffice'], capture_output=True)
    time.sleep(2)

    # Remove leftover .docx
    if os.path.exists(DOCX_PATH):
        os.remove(DOCX_PATH)
        print(f'Removed leftover .docx: {DOCX_PATH}')


def create_initial():
    # Step 1: Create base document as .docx
    create_base_document()

    # Step 2: Open in LO via UNO, save as .odt, kill LO
    convert_to_odt()

    # Step 3: Relaunch LO with the .odt file for GUI-ready state
    launch_gui(f'libreoffice --writer "{ODT_PATH}"', delay_sec=3.0)
    print(f'Initial file created: {ODT_PATH}')
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
