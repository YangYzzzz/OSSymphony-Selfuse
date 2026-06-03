"""
Reward Script: Quick Reference Card formatting task
Task ID: writer_mktg_058
Domain: libreoffice_writer
Scoring:
  Component 1 (0.15): Title 'Quick Reference Card' — 18pt bold, centered
  Component 2 (0.30): 2x4 table (4 rows x 2 cols) with all 8 fact labels + values present
  Component 3 (0.15): Table cell formatting — labels 10pt bold, values 14pt bold
  Component 4 (0.15): Alternating cell backgrounds — left=E3F2FD, right=FFFFFF
  Component 5 (0.10): At least 1 manual page break between page 1 and page 2
  Component 6 (0.15): Page 2 brand messaging — 'Brand Messaging' heading, 'Elevator Pitch'
                       subheading with italic 14pt text, 'Mission Statement' box,
                       and 3 numbered key messages with bold lead-in phrases
Total: 1.0
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_mktg_058'
FILE_PATH = f'{WORKDIR}/Desktop/company_quick_reference.docx'

# Expected fact table contents: label -> value
EXPECTED_FACTS = {
    'Founded': '2018',
    'HQ': 'San Francisco',
    'Employees': '310',
    'Revenue': '$67M ARR',
    'Customers': '2,800+',
    'Products': '3',
    'Markets': '15 countries',
    'Awards': '12 industry awards',
}

# Expected key message lead-ins
KEY_MESSAGE_LEADS = ['Simplicity at Scale', 'Proven ROI', 'Partnership, Not Just Software']


def check_title(doc):
    """Return True if title 'Quick Reference Card' is 18pt, bold, centered."""
    for para in doc.paragraphs:
        if 'Quick Reference Card' not in para.text.strip():
            continue
        alignment = para.paragraph_format.alignment
        if alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            continue
        for run in para.runs:
            sz = run.font.size
            sz_pt = sz.pt if sz else None
            if run.bold is True and sz_pt is not None and abs(sz_pt - 18.0) < 0.5:
                return True
    return False


def find_fact_table(doc):
    """Return the 4-row x 2-col fact table, or None."""
    for table in doc.tables:
        if len(table.rows) == 4 and len(table.columns) == 2:
            return table
    return None


def check_fact_table_contents(fact_table):
    """Return True if all 8 expected fact labels and values are present in the table."""
    cell_texts = []
    for row in fact_table.rows:
        for cell in row.cells:
            cell_texts.append(cell.text.strip())
    all_text = ' '.join(cell_texts)
    missing_labels = [l for l in EXPECTED_FACTS if l not in all_text]
    missing_values = [v for v in EXPECTED_FACTS.values() if v not in all_text]
    return missing_labels, missing_values


def check_cell_formatting(fact_table):
    """Return count of cells with correct label (10pt bold) + value (14pt bold) formatting."""
    ok_count = 0
    total_cells = 0
    for row in fact_table.rows:
        for cell in row.cells:
            total_cells += 1
            cell_runs = []
            for p in cell.paragraphs:
                cell_runs.extend(p.runs)
            if len(cell_runs) >= 2:
                label_run = cell_runs[0]
                value_run = cell_runs[1]
                label_sz = label_run.font.size
                value_sz = value_run.font.size
                label_pt = label_sz.pt if label_sz else None
                value_pt = value_sz.pt if value_sz else None
                label_bold = label_run.bold is True
                value_bold = value_run.bold is True
                if (label_pt is not None and abs(label_pt - 10.0) < 0.5 and label_bold and
                        value_pt is not None and abs(value_pt - 14.0) < 0.5 and value_bold):
                    ok_count += 1
    return ok_count, total_cells


def check_cell_backgrounds(fact_table):
    """Return count of cells with correct alternating backgrounds."""
    ok_count = 0
    total_bg_cells = 0
    for ri, row in enumerate(fact_table.rows):
        for ci, cell in enumerate(row.cells):
            total_bg_cells += 1
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            fill = None
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill:
                        fill = fill.upper()
            expected_fill = 'E3F2FD' if ci == 0 else 'FFFFFF'
            if fill == expected_fill:
                ok_count += 1
    return ok_count, total_bg_cells


def count_page_breaks(doc):
    """Return number of manual page breaks in the document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for br in run.element.findall('.//w:br', ns):
                br_type = br.attrib.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'
                )
                if br_type == 'page':
                    count += 1
    return count


def check_brand_messaging(doc):
    """Return dict of brand messaging sub-check results."""
    results = {
        'brand_messaging_bold': False,
        'elevator_pitch_heading': False,
        'elevator_pitch_italic14pt': False,
        'mission_statement_heading': False,
        'mission_box': False,
        'key_messages_count': 0,
    }

    for para in doc.paragraphs:
        text = para.text.strip()

        # 'Brand Messaging' bold heading
        if text == 'Brand Messaging':
            for run in para.runs:
                if run.bold is True:
                    results['brand_messaging_bold'] = True
                    break

        # 'Elevator Pitch' subheading present
        if text == 'Elevator Pitch':
            results['elevator_pitch_heading'] = True

        # Elevator pitch text in italic 14pt
        if 'We help growing businesses' in text or 'streamline their operations' in text:
            for run in para.runs:
                sz = run.font.size
                sz_pt = sz.pt if sz else None
                if run.italic is True and sz_pt is not None and abs(sz_pt - 14.0) < 0.5:
                    results['elevator_pitch_italic14pt'] = True
                    break

        # 'Mission Statement' heading
        if text == 'Mission Statement':
            results['mission_statement_heading'] = True

        # Key messages with bold lead-ins
        for lead in KEY_MESSAGE_LEADS:
            if lead in text:
                for run in para.runs:
                    if lead in run.text and run.bold is True:
                        results['key_messages_count'] += 1
                        break

    # Mission statement in a single-cell table (bordered box)
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 1:
            cell_text = table.cell(0, 0).text.strip()
            if 'empower businesses' in cell_text or 'work smarter' in cell_text:
                results['mission_box'] = True
                break

    return results


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # Precondition: file must be loadable
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Title 'Quick Reference Card' — 18pt, bold, centered (0.15)
    # -----------------------------------------------------------------------
    try:
        if check_title(doc):
            print("PASS: Component 1 — Title 'Quick Reference Card' found: 18pt, bold, centered (0.15 pts)")
            total_score += 0.15
        else:
            print("FAIL: Component 1 — Title not found with 18pt bold centered formatting")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: 2x4 table (4 rows x 2 cols) with all 8 fact labels + values (0.30)
    # -----------------------------------------------------------------------
    fact_table = None
    try:
        fact_table = find_fact_table(doc)
        if fact_table is None:
            print(f"FAIL: Component 2 — No 4-row x 2-col fact table found (tables: {len(doc.tables)})")
        else:
            missing_labels, missing_values = check_fact_table_contents(fact_table)
            if not missing_labels and not missing_values:
                print("PASS: Component 2 — 4x2 fact table with all 8 labels and values present (0.30 pts)")
                total_score += 0.30
            else:
                print(f"FAIL: Component 2 — Missing labels: {missing_labels}, missing values: {missing_values}")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Cell formatting — label 10pt bold, value 14pt bold (0.15)
    # -----------------------------------------------------------------------
    try:
        if fact_table is None:
            print("FAIL: Component 3 — No fact table to check formatting")
        else:
            ok_count, total_cells = check_cell_formatting(fact_table)
            if ok_count == total_cells and total_cells == 8:
                print(f"PASS: Component 3 — All {total_cells} cells have label 10pt bold + value 14pt bold (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 3 — Only {ok_count}/{total_cells} cells have correct formatting")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # -----------------------------------------------------------------------
    # Component 4: Alternating cell backgrounds E3F2FD / FFFFFF (0.15)
    # -----------------------------------------------------------------------
    try:
        if fact_table is None:
            print("FAIL: Component 4 — No fact table to check backgrounds")
        else:
            bg_ok, bg_total = check_cell_backgrounds(fact_table)
            if bg_ok == bg_total and bg_total == 8:
                print(f"PASS: Component 4 — All {bg_total} cells have correct backgrounds (0.15 pts)")
                total_score += 0.15
            else:
                print(f"FAIL: Component 4 — Only {bg_ok}/{bg_total} cells have correct backgrounds")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # -----------------------------------------------------------------------
    # Component 5: At least 1 manual page break (0.10)
    # -----------------------------------------------------------------------
    try:
        pb_count = count_page_breaks(doc)
        if pb_count >= 1:
            print(f"PASS: Component 5 — {pb_count} manual page break(s) found (0.10 pts)")
            total_score += 0.10
        else:
            print("FAIL: Component 5 — No manual page breaks found (expected at least 1)")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    # -----------------------------------------------------------------------
    # Component 6: Page 2 brand messaging structure (0.15)
    # -----------------------------------------------------------------------
    try:
        brand = check_brand_messaging(doc)
        sub_checks_passed = (
            int(brand['brand_messaging_bold']) +
            int(brand['elevator_pitch_heading']) +
            int(brand['elevator_pitch_italic14pt']) +
            int(brand['mission_statement_heading']) +
            int(brand['mission_box']) +
            int(brand['key_messages_count'] >= 3)
        )
        details = (f"brand_messaging={brand['brand_messaging_bold']}, "
                   f"elevator_heading={brand['elevator_pitch_heading']}, "
                   f"elevator_italic14pt={brand['elevator_pitch_italic14pt']}, "
                   f"mission_heading={brand['mission_statement_heading']}, "
                   f"mission_box={brand['mission_box']}, "
                   f"key_messages={brand['key_messages_count']}/3")

        if sub_checks_passed == 6:
            print(f"PASS: Component 6 — All brand messaging elements present (0.15 pts)")
            print(f"  {details}")
            total_score += 0.15
        elif sub_checks_passed >= 3:
            print(f"PARTIAL: Component 6 — {sub_checks_passed}/6 sub-checks passed (0.08 pts)")
            print(f"  {details}")
            total_score += 0.08
        else:
            print(f"FAIL: Component 6 — Only {sub_checks_passed}/6 sub-checks passed")
            print(f"  {details}")
    except Exception as e:
        print(f"ERROR: Component 6 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.2f}/1.0")
    print(f"REWARD: {final_score:.1f}")
    return final_score


if not os.path.exists(FILE_PATH):
    print(f"File not found: {FILE_PATH}")
    print("REWARD: 0.0")
else:
    verify_task(FILE_PATH)
