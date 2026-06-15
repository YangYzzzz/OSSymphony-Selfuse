"""
Initial Setup: Company Quick Reference Card - unformatted plain text version
Task ID: writer_mktg_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user/Desktop'  # VM path — file goes to ~/Desktop/
TASK_ID = 'company_quick_reference'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    # Ensure Desktop directory exists
    os.makedirs(WORKDIR, exist_ok=True)

    doc = Document()

    # Remove default styles that might add unexpected formatting
    # Use Normal style throughout — all plain 12pt text

    # --- Company Facts Section ---
    p = doc.add_paragraph()
    run = p.add_run("Company Quick Reference Card")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Company Facts")
    run.font.size = Pt(12)

    facts = [
        ("Founded", "2018"),
        ("HQ", "San Francisco"),
        ("Employees", "310"),
        ("Revenue", "$67M ARR"),
        ("Customers", "2,800+"),
        ("Products", "3"),
        ("Markets", "15 countries"),
        ("Awards", "12 industry awards"),
    ]

    for label, value in facts:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: {value}")
        run.font.size = Pt(12)

    # --- Elevator Pitch ---
    p = doc.add_paragraph()
    run = p.add_run("Elevator Pitch")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "We help growing businesses streamline their operations through intelligent automation "
        "that connects every part of their workflow. Our platform reduces manual work by 80%, "
        "letting teams focus on what matters most: delighting customers and driving growth."
    )
    run.font.size = Pt(12)

    # --- Mission Statement ---
    p = doc.add_paragraph()
    run = p.add_run("Mission Statement")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        "To empower businesses of every size with the tools they need to work smarter, "
        "grow faster, and build lasting customer relationships."
    )
    run.font.size = Pt(12)

    # --- Three Key Messages ---
    p = doc.add_paragraph()
    run = p.add_run("Key Messages")
    run.font.size = Pt(12)

    key_messages = [
        (
            "Simplicity at Scale:",
            " Our platform is designed to be intuitive from day one, yet powerful enough "
            "to handle enterprise-level complexity. Customers are up and running in under 48 hours "
            "with zero infrastructure investment."
        ),
        (
            "Proven ROI:",
            " Across our 2,800+ customers in 15 markets, the average return on investment "
            "is 340% within the first year. Our clients consistently report saving 15+ hours "
            "per week per team, translating directly to bottom-line impact."
        ),
        (
            "Partnership, Not Just Software:",
            " We assign a dedicated success manager to every account. From onboarding to "
            "ongoing optimization, our team acts as an extension of yours — ensuring you always "
            "get maximum value from the platform."
        ),
    ]

    for lead_in, body in key_messages:
        p = doc.add_paragraph()
        run = p.add_run(lead_in + body)
        run.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup — open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
