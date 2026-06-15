"""
Reward Script: Change body text font to Times New Roman 12pt in legal brief
Task ID: writer_legal_002
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): Normal style font name is Times New Roman
  Component 2 (0.3): Normal style font size is 12pt
  Component 3 (0.25): Body paragraph runs reflect Times New Roman 12pt
  Component 4 (0.15): Heading styles retain Liberation Sans (not overwritten)
"""

import os
from docx import Document
from docx.shared import Pt

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_002'


def persist_app_state(domain):
    """Save any unsaved LibreOffice edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for {}".format(domain))
    except Exception as e:
        print("PERSIST_WARN: save hook failed: {}".format(e))


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file {}: {}".format(file_path, e))
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Normal style font name is Times New Roman (0.3 points)
    # Initial: Liberation Sans -> Golden: Times New Roman
    try:
        normal_style = doc.styles['Normal']
        font_name = normal_style.font.name
        if font_name == 'Times New Roman':
            print("PASS: Component 1 -- Normal style font name is '{}' (0.3 pts)".format(font_name))
            total_score += 0.3
        else:
            print("FAIL: Component 1 -- Expected Normal style font 'Times New Roman', found '{}'".format(font_name))
    except Exception as e:
        print("ERROR: Component 1 -- {}".format(e))

    # Component 2: Normal style font size is 12pt (0.3 points)
    # Initial: 11pt -> Golden: 12pt
    try:
        normal_style = doc.styles['Normal']
        font_size = normal_style.font.size
        if font_size is not None and abs(font_size.pt - 12.0) < 0.1:
            print("PASS: Component 2 -- Normal style font size is {}pt (0.3 pts)".format(font_size.pt))
            total_score += 0.3
        else:
            size_str = "{}pt".format(font_size.pt) if font_size else "None"
            print("FAIL: Component 2 -- Expected Normal style font size 12pt, found {}".format(size_str))
    except Exception as e:
        print("ERROR: Component 2 -- {}".format(e))

    # Component 3: Body paragraph runs reflect Times New Roman 12pt (0.25 points)
    # Check that actual body text runs have the correct font
    try:
        body_runs_checked = 0
        body_runs_correct = 0
        for para in doc.paragraphs:
            # Skip headings
            if para.style and 'Heading' in para.style.name:
                continue
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            for run in para.runs:
                if not run.text.strip():
                    continue
                body_runs_checked += 1
                rname = run.font.name
                rsize = run.font.size
                # Font name must be Times New Roman (or None = inherited from style)
                name_ok = (rname == 'Times New Roman') or (rname is None)
                # Font size must be 12pt (or None = inherited from style)
                size_ok = (rsize is None) or (rsize is not None and abs(rsize.pt - 12.0) < 0.1)
                if name_ok and size_ok:
                    body_runs_correct += 1

        if body_runs_checked == 0:
            print("FAIL: Component 3 -- No body runs found to check")
        else:
            ratio = body_runs_correct / body_runs_checked
            if ratio >= 0.9:
                print("PASS: Component 3 -- {}/{} body runs have Times New Roman 12pt (0.25 pts)".format(
                    body_runs_correct, body_runs_checked))
                total_score += 0.25
            else:
                print("FAIL: Component 3 -- Only {}/{} body runs have correct font ({:.0%})".format(
                    body_runs_correct, body_runs_checked, ratio))
    except Exception as e:
        print("ERROR: Component 3 -- {}".format(e))

    # Component 4: Heading paragraphs retain Liberation Sans (0.15 points)
    # Headings must NOT be changed to Times New Roman
    try:
        heading_runs_checked = 0
        heading_runs_retained = 0
        for para in doc.paragraphs:
            if not (para.style and 'Heading' in para.style.name):
                continue
            for run in para.runs:
                if not run.text.strip():
                    continue
                heading_runs_checked += 1
                rname = run.font.name
                # Headings should still be Liberation Sans (not Times New Roman)
                if rname == 'Liberation Sans':
                    heading_runs_retained += 1
                elif rname is None:
                    # Inherited from heading style, which should not be TNR
                    heading_runs_retained += 1

        if heading_runs_checked == 0:
            print("FAIL: Component 4 -- No heading runs found to check")
        else:
            ratio = heading_runs_retained / heading_runs_checked
            # This component only awards points if BOTH headings are retained AND
            # the Normal style was changed (i.e., Component 1 passed).
            # This ensures it doesn't award points on initial_env where headings
            # are also Liberation Sans but task hasn't been done.
            normal_font = doc.styles['Normal'].font.name
            if normal_font == 'Times New Roman' and ratio >= 0.9:
                print("PASS: Component 4 -- {}/{} heading runs retain Liberation Sans while body changed (0.15 pts)".format(
                    heading_runs_retained, heading_runs_checked))
                total_score += 0.15
            elif normal_font != 'Times New Roman':
                print("FAIL: Component 4 -- Normal style not yet Times New Roman, heading check deferred")
            else:
                print("FAIL: Component 4 -- Only {}/{} heading runs retain original font ({:.0%})".format(
                    heading_runs_retained, heading_runs_checked, ratio))
    except Exception as e:
        print("ERROR: Component 4 -- {}".format(e))

    final_score = min(total_score, 1.0)
    print("\nScore: {}/1.0".format(total_score))
    print("REWARD: {}".format(final_score))
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = '{}/{}.docx'.format(WORKDIR, TASK_ID)
if not os.path.exists(file_path):
    print("File not found: {}".format(file_path))
    print("REWARD: 0.0")
else:
    verify_task(file_path)
