"""
Initial Setup: Legal brief with body text in Liberation Sans 11pt
Task ID: writer_legal_002
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_002'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def set_run_font(run, font_name, font_size_pt, bold=False, italic=False):
    """Helper to set font properties on a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    run.italic = italic
    # Set eastAsia font too for completeness
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_body_paragraph(doc, text, alignment=None, space_after_pt=6):
    """Add a body paragraph with Liberation Sans 11pt."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, 'Liberation Sans', 11)
    if alignment:
        para.paragraph_format.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after_pt)
    return para


def add_heading_styled(doc, text, level=1):
    """Add a heading using built-in heading style."""
    heading = doc.add_heading(text, level=level)
    # Headings use their own style (Heading 1, Heading 2, etc.)
    # We explicitly set the heading font to Liberation Sans Bold at larger sizes
    # to distinguish from body text
    for run in heading.runs:
        if level == 1:
            set_run_font(run, 'Liberation Sans', 16, bold=True)
        elif level == 2:
            set_run_font(run, 'Liberation Sans', 14, bold=True)
        elif level == 3:
            set_run_font(run, 'Liberation Sans', 12, bold=True, italic=True)
    return heading


def create_initial():
    doc = Document()

    # Set the default paragraph style font to Liberation Sans 11pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Liberation Sans'
    font.size = Pt(11)

    # --- Page Setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Title Block ---
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run('IN THE CIRCUIT COURT OF COOK COUNTY, ILLINOIS')
    set_run_font(run, 'Liberation Sans', 11, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run('COUNTY DEPARTMENT, LAW DIVISION')
    set_run_font(run, 'Liberation Sans', 11, bold=True)

    # Parties
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, 'GREENFIELD TECHNOLOGIES, INC.,')
    p = add_body_paragraph(doc, '\tPlaintiff,')
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, 'v.\t\t\t\tCase No. 2025-L-004837')
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, 'SUMMIT RIDGE CAPITAL PARTNERS, LLC,')
    p = add_body_paragraph(doc, '\tDefendant.')
    add_body_paragraph(doc, '')

    # Document title
    brief_title = doc.add_paragraph()
    brief_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    brief_title.paragraph_format.space_after = Pt(12)
    run = brief_title.add_run("PLAINTIFF'S MEMORANDUM IN SUPPORT OF MOTION FOR SUMMARY JUDGMENT")
    set_run_font(run, 'Liberation Sans', 11, bold=True)

    # --- Section I ---
    add_heading_styled(doc, 'I. INTRODUCTION', level=1)

    add_body_paragraph(doc,
        'Plaintiff Greenfield Technologies, Inc. ("Greenfield") respectfully submits this '
        'memorandum in support of its motion for summary judgment pursuant to 735 ILCS 5/2-1005. '
        'The undisputed material facts establish that Defendant Summit Ridge Capital Partners, LLC '
        '("Summit Ridge") breached the parties\' Master Services Agreement dated March 15, 2023 '
        '(the "Agreement") by failing to remit payment for professional consulting services '
        'rendered between April 2023 and September 2024.')

    add_body_paragraph(doc,
        'As demonstrated below, there is no genuine issue of material fact regarding the existence '
        'of the Agreement, Greenfield\'s full performance thereunder, Summit Ridge\'s failure to pay '
        'the invoiced amounts totaling $487,230.00, and the resulting damages. Accordingly, '
        'Greenfield is entitled to judgment as a matter of law.')

    # --- Section II ---
    add_heading_styled(doc, 'II. STATEMENT OF UNDISPUTED MATERIAL FACTS', level=1)

    add_body_paragraph(doc,
        '1. On March 15, 2023, Greenfield and Summit Ridge executed a Master Services Agreement '
        'under which Greenfield agreed to provide technology consulting and software integration '
        'services to Summit Ridge. (Exhibit A, Agreement.)')

    add_body_paragraph(doc,
        '2. Section 4.1 of the Agreement required Summit Ridge to pay all invoices within thirty '
        '(30) days of receipt. (Ex. A, \u00a7 4.1.)')

    add_body_paragraph(doc,
        '3. Between April 1, 2023 and September 30, 2024, Greenfield provided consulting services '
        'in accordance with the Statement of Work attached as Exhibit B to the Agreement. '
        '(Affidavit of Rachel Torres, \u00b6\u00b6 5-8; Ex. B.)')

    add_body_paragraph(doc,
        '4. Greenfield issued twelve (12) monthly invoices to Summit Ridge, each detailing the '
        'services performed and fees owed. The total amount invoiced was $487,230.00. '
        '(Torres Aff., \u00b6 9; Ex. C, Invoices.)')

    add_body_paragraph(doc,
        '5. Summit Ridge made partial payments totaling $112,500.00 between May 2023 and '
        'August 2023, leaving an outstanding balance of $374,730.00. (Torres Aff., \u00b6 11; '
        'Ex. D, Payment Ledger.)')

    add_body_paragraph(doc,
        '6. On October 15, 2024, Greenfield sent a formal demand letter to Summit Ridge via '
        'certified mail, demanding payment of the outstanding balance. (Ex. E, Demand Letter; '
        'Ex. F, Certified Mail Receipt.)')

    add_body_paragraph(doc,
        '7. Summit Ridge did not respond to the demand letter and has made no further payments '
        'since August 2023. (Torres Aff., \u00b6 14.)')

    # --- Section III ---
    add_heading_styled(doc, 'III. LEGAL STANDARD', level=1)

    add_body_paragraph(doc,
        'Summary judgment is appropriate when "the pleadings, depositions, and admissions on file, '
        'together with the affidavits, if any, show that there is no genuine issue as to any '
        'material fact and that the moving party is entitled to a judgment as a matter of law." '
        '735 ILCS 5/2-1005(c). The purpose of summary judgment is not to try an issue of fact '
        'but rather to determine whether one exists. Adams v. Northern Illinois Gas Co., '
        '211 Ill. 2d 32, 43 (2004).')

    add_body_paragraph(doc,
        'The moving party bears the initial burden of production, which may be met by affirmatively '
        'demonstrating the absence of a genuine issue of material fact or by establishing that the '
        'nonmoving party\'s case is insufficient as a matter of law. Nettleton v. Stogsdill, '
        '387 Ill. App. 3d 743, 756 (4th Dist. 2008).')

    # --- Section IV ---
    add_heading_styled(doc, 'IV. ARGUMENT', level=1)

    add_heading_styled(doc, 'A. Greenfield Has Established All Elements of Breach of Contract', level=2)

    add_body_paragraph(doc,
        'Under Illinois law, to prevail on a breach of contract claim, a plaintiff must demonstrate: '
        '(1) the existence of a valid and enforceable contract; (2) substantial performance by the '
        'plaintiff; (3) a breach by the defendant; and (4) resultant damages. Henderson-Smith & '
        'Associates, Inc. v. Nahamani Family Service Center, Inc., 323 Ill. App. 3d 15, 27 '
        '(1st Dist. 2001).')

    add_heading_styled(doc, '1. A Valid Contract Exists Between the Parties', level=3)

    add_body_paragraph(doc,
        'The Master Services Agreement satisfies all requirements of a valid contract under '
        'Illinois law. The Agreement was executed by authorized representatives of both parties, '
        'contains definite and certain terms regarding the scope of services and payment obligations, '
        'and is supported by adequate consideration. (Ex. A.) Summit Ridge has not disputed the '
        'validity or enforceability of the Agreement at any point during this litigation.')

    add_heading_styled(doc, '2. Greenfield Substantially Performed Under the Agreement', level=3)

    add_body_paragraph(doc,
        'The uncontroverted evidence demonstrates that Greenfield performed all consulting services '
        'required under the Statement of Work. Ms. Torres\'s affidavit details the specific services '
        'provided during each billing period, and Summit Ridge has produced no evidence suggesting '
        'that Greenfield\'s performance was deficient in any respect. (Torres Aff., \u00b6\u00b6 5-8.)')

    add_heading_styled(doc, '3. Summit Ridge Breached the Agreement by Failing to Pay', level=3)

    add_body_paragraph(doc,
        'Summit Ridge\'s obligation to pay invoices within thirty days is unambiguous. The payment '
        'ledger confirms that Summit Ridge ceased making payments after August 2023, leaving '
        '$374,730.00 in unpaid invoices. This failure to pay constitutes a material breach of '
        'Section 4.1 of the Agreement. (Torres Aff., \u00b6\u00b6 9-11; Ex. D.)')

    add_heading_styled(doc, '4. Greenfield Has Suffered Damages', level=3)

    add_body_paragraph(doc,
        'As a direct and proximate result of Summit Ridge\'s breach, Greenfield has been damaged in '
        'the amount of $374,730.00 representing unpaid invoices, plus pre-judgment interest at the '
        'statutory rate of 5% per annum pursuant to 815 ILCS 205/2 from the date each invoice '
        'became past due.')

    add_heading_styled(doc, 'B. Summit Ridge Cannot Raise a Genuine Issue of Material Fact', level=2)

    add_body_paragraph(doc,
        'In its Answer, Summit Ridge raised only general denials and vague allegations that '
        'Greenfield\'s services did not meet "industry standards." However, Summit Ridge has '
        'produced no expert testimony, documentary evidence, or affidavit to support this defense. '
        'Mere allegations or denials in the pleadings are insufficient to create a genuine issue '
        'of material fact. Robidoux v. Oliphant, 201 Ill. 2d 324, 335 (2002).')

    add_body_paragraph(doc,
        'Moreover, Summit Ridge\'s partial payments totaling $112,500.00 constitute an admission '
        'that the services were rendered and accepted. A party who accepts the benefits of a '
        'contract cannot subsequently claim non-performance without specific evidence of deficiency. '
        'Martindell v. Lake Shore National Bank, 15 Ill. 2d 272, 283 (1958).')

    # --- Section V ---
    add_heading_styled(doc, 'V. CONCLUSION', level=1)

    add_body_paragraph(doc,
        'For the foregoing reasons, Plaintiff Greenfield Technologies, Inc. respectfully requests '
        'that this Court grant summary judgment in its favor and against Defendant Summit Ridge '
        'Capital Partners, LLC on Count I of the Complaint, and award Greenfield damages in the '
        'amount of $374,730.00 plus pre-judgment interest, costs, and such other relief as this '
        'Court deems just and proper.')

    add_body_paragraph(doc, '')

    # Signature block
    add_body_paragraph(doc, 'Respectfully submitted,')
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, '')

    sig = doc.add_paragraph()
    run = sig.add_run('________________________________')
    set_run_font(run, 'Liberation Sans', 11)

    add_body_paragraph(doc, 'Elena M. Vasquez, Esq.')
    add_body_paragraph(doc, 'VASQUEZ & CASTELLANO LLP')
    add_body_paragraph(doc, '180 North LaSalle Street, Suite 3200')
    add_body_paragraph(doc, 'Chicago, Illinois 60601')
    add_body_paragraph(doc, 'Tel: (312) 555-0147')
    add_body_paragraph(doc, 'evasquez@vasquezcastellano.com')
    add_body_paragraph(doc, '')
    add_body_paragraph(doc, 'Attorneys for Plaintiff')
    add_body_paragraph(doc, 'Greenfield Technologies, Inc.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
