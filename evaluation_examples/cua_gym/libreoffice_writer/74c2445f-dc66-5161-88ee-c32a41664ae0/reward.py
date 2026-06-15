"""
FINAL REWARD SCRIPT - SUCCESS
Task: Could you set the 2 in "H2SO4" as a subscript?
Generated: 2025-10-14 07:16:26
Status: success
Model: azure-o3
Total Steps: 3
"""

from docx import Document
import os

def verify_h2so4_subscript(file_path: str) -> float:
    """Reward script for the task:
    "Could you set the 2 in 'H2SO4' as a subscript?"

    Scoring (progressive):
    0.2  – The text "H2SO4" is present in the document (case-insensitive).
    0.8  – In a paragraph containing that text, the character "2" is formatted as subscript
            (run.font.subscript == True or WD_BOOL TRUE).
    -----
    1.0  – Both conditions satisfied.
    """

    print(f"Verifying file: {file_path}")
    total_score = 0.0
    max_score = 1.0

    # -------------- Pre-checks --------------
    if not os.path.exists(file_path):
        print("✗ File does not exist")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"✗ Unable to load DOCX: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -------------- Requirement 1: presence of H2SO4 --------------
    all_text = " ".join(p.text for p in doc.paragraphs)
    if "h2so4" in all_text.lower():
        print("✓ Found text 'H2SO4' in document (0.2 points)")
        total_score += 0.2
    else:
        print("✗ Text 'H2SO4' not found in document")

    # -------------- Requirement 2: '2' formatted as subscript --------------
    subscript_ok = False
    for para in doc.paragraphs:
        if "h2so4" in para.text.lower():
            for run in para.runs:
                if "2" in run.text:
                    sub_val = run.font.subscript  # True / False / None / WD_BOOL
                    if sub_val is True or (hasattr(sub_val, "value") and sub_val.value == 1):
                        subscript_ok = True
                        print(f"✓ Run containing '2' (text='{run.text}') is subscript (0.8 points)")
                        break
            if subscript_ok:
                break

    if subscript_ok:
        total_score += 0.8
    else:
        print("✗ Did not detect '2' formatted as subscript in context of 'H2SO4'")

    # -------------- Final score --------------
    total_score = min(total_score, max_score)
    print(f"Total Score: {total_score}/{max_score}")
    print(f"REWARD: {total_score}")
    return total_score


if __name__ == "__main__":
    # Default path for the autograder environment
    path = "/home/user/could_you_set_the_2_in_h2so4_as_a_subscript.docx"
    verify_h2so4_subscript(path)

