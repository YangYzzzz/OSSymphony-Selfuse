"""
Reward Script: Create labels with company name and QR code placeholder
Task ID: writer_lec_061
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): Outer label grid table exists (Avery 5160 layout ~10 rows x 3 cols)
  Component 2 (0.30): Labels contain 'TechVentures LLC' text
  Component 3 (0.30): QR code placeholder cells with visible borders
  Component 4 (0.15): QR placeholder is approximately 1.5 cm (850 dxa)
"""

import os
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_lec_061'


def persist_app_state(domain):
    """Try to save any unsaved changes via Ctrl+S."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for %s" % domain)
    except Exception as e:
        print("PERSIST_WARN: save hook failed: %s" % e)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        import xml.etree.ElementTree as ET
    except ImportError as e:
        print("CRITICAL: Missing library: %s" % e)
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print("CRITICAL: Cannot load file %s: %s" % (file_path, e))
        print("REWARD: 0.0")
        return 0.0

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Component 1: Outer label grid table exists (0.25 points)
    # Avery 5160 is 10 rows x 3 cols = 30 labels
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 1 -- No tables found in document")
        else:
            t = doc.tables[0]
            nrows = len(t.rows)
            ncols = len(t.columns)
            # Avery 5160: 10 rows x 3 cols. Accept reasonable range.
            if ncols == 3 and nrows >= 5:
                print("PASS: Component 1 -- Label grid table found: %d rows x %d cols (0.25 pts)" % (nrows, ncols))
                total_score += 0.25
            elif ncols >= 2 and nrows >= 3:
                # Partial: some grid structure but not standard 5160
                print("PARTIAL: Component 1 -- Table found but non-standard: %d rows x %d cols (0.10 pts)" % (nrows, ncols))
                total_score += 0.10
            else:
                print("FAIL: Component 1 -- Table too small: %d rows x %d cols" % (nrows, ncols))
    except Exception as e:
        print("ERROR: Component 1 -- %s" % e)

    # Component 2: Labels contain 'TechVentures LLC' text (0.30 points)
    # Check how many outer cells contain the company name
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 2 -- No tables to check")
        else:
            t = doc.tables[0]
            total_cells = 0
            cells_with_tv = 0
            for row in t.rows:
                for cell in row.cells:
                    total_cells += 1
                    cell_xml = ET.tostring(cell._element, encoding="unicode")
                    if "TechVentures" in cell_xml:
                        cells_with_tv += 1

            if total_cells == 0:
                print("FAIL: Component 2 -- No cells in table")
            else:
                ratio = cells_with_tv / total_cells
                if ratio >= 0.8:
                    print("PASS: Component 2 -- %d/%d cells contain 'TechVentures LLC' (0.30 pts)" % (cells_with_tv, total_cells))
                    total_score += 0.30
                elif ratio >= 0.4:
                    print("PARTIAL: Component 2 -- %d/%d cells contain 'TechVentures LLC' (0.15 pts)" % (cells_with_tv, total_cells))
                    total_score += 0.15
                else:
                    print("FAIL: Component 2 -- Only %d/%d cells contain 'TechVentures LLC'" % (cells_with_tv, total_cells))
    except Exception as e:
        print("ERROR: Component 2 -- %s" % e)

    # Component 3: QR code placeholder cells with visible borders (0.30 points)
    # Each label should have a nested table; the second cell of that nested table
    # should have visible (single) borders on all 4 sides.
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 3 -- No tables to check")
        else:
            t = doc.tables[0]
            nested_tbls = t._element.findall(".//w:tbl", ns)
            if len(nested_tbls) == 0:
                print("FAIL: Component 3 -- No nested tables (no QR placeholders)")
            else:
                bordered_count = 0
                for ntbl in nested_tbls:
                    tcs = ntbl.findall(".//w:tc", ns)
                    if len(tcs) < 2:
                        continue
                    # Check last cell for borders (QR placeholder)
                    qr_tc = tcs[-1]
                    tcpr = qr_tc.find("w:tcPr", ns)
                    if tcpr is None:
                        continue
                    borders = tcpr.find("w:tcBorders", ns)
                    if borders is None:
                        continue
                    # Check for visible borders on all sides
                    border_sides = 0
                    for child in borders:
                        tag = child.tag.split("}")[-1]
                        val = child.get(qn("w:val"))
                        if val and val != "none" and val != "nil":
                            border_sides += 1
                    if border_sides >= 4:
                        bordered_count += 1

                total_labels = len(t.rows) * len(t.columns)
                if bordered_count >= total_labels * 0.8:
                    print("PASS: Component 3 -- %d/%d labels have bordered QR placeholder (0.30 pts)" % (bordered_count, total_labels))
                    total_score += 0.30
                elif bordered_count >= total_labels * 0.3:
                    print("PARTIAL: Component 3 -- %d/%d labels have bordered QR placeholder (0.15 pts)" % (bordered_count, total_labels))
                    total_score += 0.15
                else:
                    print("FAIL: Component 3 -- Only %d/%d labels have bordered QR placeholder" % (bordered_count, total_labels))
    except Exception as e:
        print("ERROR: Component 3 -- %s" % e)

    # Component 4: QR placeholder is approximately 1.5 cm (0.15 points)
    # 1.5 cm = ~850 dxa (twips). Accept 600-1100 dxa range.
    try:
        if len(doc.tables) == 0:
            print("FAIL: Component 4 -- No tables to check")
        else:
            t = doc.tables[0]
            nested_tbls = t._element.findall(".//w:tbl", ns)
            if len(nested_tbls) == 0:
                print("FAIL: Component 4 -- No nested tables")
            else:
                correct_size_count = 0
                checked = 0
                for ntbl in nested_tbls:
                    tcs = ntbl.findall(".//w:tc", ns)
                    if len(tcs) < 2:
                        continue
                    qr_tc = tcs[-1]
                    tcpr = qr_tc.find("w:tcPr", ns)
                    if tcpr is None:
                        continue
                    # Get cell width from tcW elements (use LAST one as override)
                    tcw_elems = tcpr.findall("w:tcW", ns)
                    if tcw_elems:
                        tcw = tcw_elems[-1]  # Last tcW is the effective one
                        w_val = tcw.get(qn("w:w"))
                        w_type = tcw.get(qn("w:type"))
                        if w_val and w_type == "dxa":
                            checked += 1
                            w_int = int(w_val)
                            # 1.5 cm ~ 850 dxa. Allow 600-1100 range.
                            if 600 <= w_int <= 1100:
                                correct_size_count += 1

                if checked == 0:
                    print("FAIL: Component 4 -- No QR cell width data found")
                elif correct_size_count >= checked * 0.8:
                    print("PASS: Component 4 -- %d/%d QR placeholders ~1.5cm width (0.15 pts)" % (correct_size_count, checked))
                    total_score += 0.15
                elif correct_size_count > 0:
                    print("PARTIAL: Component 4 -- %d/%d QR placeholders ~1.5cm width (0.07 pts)" % (correct_size_count, checked))
                    total_score += 0.07
                else:
                    print("FAIL: Component 4 -- 0/%d QR placeholders have ~1.5cm width" % checked)
    except Exception as e:
        print("ERROR: Component 4 -- %s" % e)

    final_score = round(min(total_score, 1.0), 2)
    print("\nScore: %.2f/1.0" % total_score)
    print("REWARD: %.1f" % final_score)
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = os.path.join(WORKDIR, TASK_ID + '.docx')
if not os.path.exists(file_path):
    print("File not found: %s" % file_path)
    print("REWARD: 0.0")
else:
    verify_task(file_path)
