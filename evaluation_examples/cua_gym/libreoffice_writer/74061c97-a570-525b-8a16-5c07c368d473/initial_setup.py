"""
Initial Setup: Create a PhD thesis document with 15 bibliography entries inserted
throughout the text, but no bibliography index at the end.
Task ID: writer_mt_077
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_077'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

# 15 bibliography entries: Author, Title, Publisher, Year
BIBLIOGRAPHY_ENTRIES = [
    {"tag": "Anderson2019", "author": "Anderson, R. M.", "title": "Computational Methods in Modern Physics", "publisher": "Cambridge University Press", "year": "2019"},
    {"tag": "Baker2020", "author": "Baker, S. L.", "title": "Quantum Field Theory and Statistical Mechanics", "publisher": "Springer", "year": "2020"},
    {"tag": "Chen2018", "author": "Chen, W.", "title": "Topological Insulators: Fundamentals and Applications", "publisher": "Academic Press", "year": "2018"},
    {"tag": "Davidson2021", "author": "Davidson, E. R.", "title": "Advanced Solid State Physics", "publisher": "Oxford University Press", "year": "2021"},
    {"tag": "Evans2017", "author": "Evans, P. G.", "title": "Synchrotron Radiation Techniques in Materials Science", "publisher": "Wiley", "year": "2017"},
    {"tag": "Fischer2022", "author": "Fischer, K. H.", "title": "Spin Glasses and Random Fields", "publisher": "World Scientific", "year": "2022"},
    {"tag": "Garcia2020", "author": "Garcia, M. A.", "title": "Nanoscale Materials Characterization", "publisher": "Elsevier", "year": "2020"},
    {"tag": "Hoffman2019", "author": "Hoffman, J. E.", "title": "Scanning Tunneling Microscopy in Surface Science", "publisher": "Springer", "year": "2019"},
    {"tag": "Ibrahim2021", "author": "Ibrahim, A. N.", "title": "Density Functional Theory: Principles and Applications", "publisher": "CRC Press", "year": "2021"},
    {"tag": "Johnson2018", "author": "Johnson, D. C.", "title": "Crystal Growth and Thin Film Deposition", "publisher": "Academic Press", "year": "2018"},
    {"tag": "Kim2023", "author": "Kim, Y. S.", "title": "Two-Dimensional Materials for Energy Applications", "publisher": "Royal Society of Chemistry", "year": "2023"},
    {"tag": "Liu2020", "author": "Liu, H.", "title": "Machine Learning Approaches to Condensed Matter Physics", "publisher": "MIT Press", "year": "2020"},
    {"tag": "Martinez2022", "author": "Martinez, C. R.", "title": "Ultrafast Spectroscopy of Quantum Materials", "publisher": "Springer Nature", "year": "2022"},
    {"tag": "Nakamura2019", "author": "Nakamura, T.", "title": "Superconductivity in Strongly Correlated Systems", "publisher": "Oxford University Press", "year": "2019"},
    {"tag": "Patel2021", "author": "Patel, V. K.", "title": "Transport Phenomena in Mesoscopic Systems", "publisher": "Cambridge University Press", "year": "2021"},
]


def add_citation_field(paragraph, entry):
    """Add a CITATION field code referencing a bibliography entry."""
    tag = entry["tag"]

    # Add field begin
    r1 = paragraph.add_run()
    fldChar1 = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fldChar1)

    # Add field instruction
    r2 = paragraph.add_run()
    instrText = r2._element.makeelement(qn('w:instrText'), {})
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' CITATION {tag} \\l 1033 '
    r2._element.append(instrText)

    # Add field separator
    r3 = paragraph.add_run()
    fldChar2 = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r3._element.append(fldChar2)

    # Add display text (cached value)
    display = f"({entry['author'].split(',')[0]}, {entry['year']})"
    r4 = paragraph.add_run(display)
    r4.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Add field end
    r5 = paragraph.add_run()
    fldChar3 = r5._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r5._element.append(fldChar3)




def create_initial():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title Page ---
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading("Quantum Transport Properties of Novel\nTwo-Dimensional Materials", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph("A Dissertation Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy in Physics")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)

    doc.add_paragraph()

    author = doc.add_paragraph("Elena M. Vasquez")
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in author.runs:
        run.font.size = Pt(16)
        run.bold = True

    doc.add_paragraph()

    dept = doc.add_paragraph("Department of Physics and Astronomy\nStanford University\nMarch 2025")
    dept.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in dept.runs:
        run.font.size = Pt(12)

    # Page break after title
    doc.add_page_break()

    # --- Abstract ---
    abstract_heading = doc.add_heading("Abstract", level=1)

    p = doc.add_paragraph(
        "This dissertation investigates the quantum transport properties of novel two-dimensional (2D) materials, "
        "with a focus on transition metal dichalcogenides (TMDs) and topological insulators. Using a combination of "
        "theoretical modeling and experimental characterization, we explore the electronic band structure, "
        "magnetotransport behavior, and spin-orbit coupling effects in atomically thin samples. "
    )
    p.add_run(
        "Our computational methods follow the framework established by "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[0])  # Anderson2019
    p.add_run(
        " with extensions to non-equilibrium conditions. The density functional theory calculations are based on "
        "the formalism detailed in "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[8])  # Ibrahim2021
    p.add_run(".")

    p2 = doc.add_paragraph(
        "We present results on quantum Hall effect measurements in high-quality MoS2 heterostructures, "
        "demonstrating fractional quantum Hall states at filling factors previously unobserved in TMD systems. "
        "Scanning tunneling microscopy measurements, performed following protocols from "
    )
    add_citation_field(p2, BIBLIOGRAPHY_ENTRIES[7])  # Hoffman2019
    p2.add_run(
        ", reveal atomic-scale defect structures that correlate with transport anomalies."
    )

    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    doc.add_heading("Chapter 1: Introduction", level=1)

    p = doc.add_paragraph(
        "The discovery of graphene in 2004 sparked an unprecedented wave of research into two-dimensional materials. "
        "Since then, the family of 2D materials has expanded dramatically to include transition metal dichalcogenides, "
        "hexagonal boron nitride, black phosphorus, and various topological materials "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[2])  # Chen2018
    p.add_run(
        ". These materials exhibit remarkable electronic properties that arise from quantum confinement "
        "and reduced dimensionality, making them ideal platforms for studying fundamental physics and developing "
        "next-generation electronic devices."
    )

    p2 = doc.add_paragraph(
        "The quantum transport properties of 2D materials are particularly fascinating because they provide "
        "direct access to phenomena that are difficult to observe in bulk systems. The integer and fractional "
        "quantum Hall effects, weak localization, and universal conductance fluctuations all manifest in "
        "qualitatively different ways in atomically thin samples. Recent advances in crystal growth techniques "
    )
    add_citation_field(p2, BIBLIOGRAPHY_ENTRIES[9])  # Johnson2018
    p2.add_run(
        " have enabled the fabrication of ultra-high quality samples with mobilities exceeding 100,000 cm2/Vs "
        "at cryogenic temperatures."
    )

    p3 = doc.add_paragraph(
        "Machine learning has emerged as a powerful tool for analyzing the complex datasets generated by "
        "transport measurements and spectroscopic techniques "
    )
    add_citation_field(p3, BIBLIOGRAPHY_ENTRIES[11])  # Liu2020
    p3.add_run(
        ". Neural network approaches have been particularly successful in identifying phase transitions "
        "and classifying topological states from transport data, opening new avenues for materials discovery."
    )

    doc.add_page_break()

    # --- Chapter 2: Theoretical Background ---
    doc.add_heading("Chapter 2: Theoretical Background", level=1)

    doc.add_heading("2.1 Electronic Band Structure of 2D Materials", level=2)

    p = doc.add_paragraph(
        "The electronic properties of two-dimensional materials are fundamentally determined by their band structure. "
        "In the tight-binding approximation, the Hamiltonian for a honeycomb lattice can be written as a 2x2 matrix "
        "in the sublattice basis, yielding the characteristic Dirac cone dispersion near the K and K' points of the "
        "Brillouin zone. For transition metal dichalcogenides, the band structure is more complex due to the "
        "contribution of d-orbitals from the metal atoms and the strong spin-orbit coupling that lifts the spin "
        "degeneracy at the valence band maximum "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[1])  # Baker2020
    p.add_run(".")

    p2 = doc.add_paragraph(
        "The density functional theory (DFT) framework provides a rigorous first-principles approach to calculating "
        "electronic band structures. In this work, we employ the generalized gradient approximation (GGA) with the "
        "Perdew-Burke-Ernzerhof (PBE) exchange-correlation functional, following the methodology outlined in "
    )
    add_citation_field(p2, BIBLIOGRAPHY_ENTRIES[8])  # Ibrahim2021
    p2.add_run(
        ". For systems with strong correlations, we supplement DFT with dynamical mean-field theory (DMFT) "
        "calculations to capture the many-body effects that are beyond the scope of standard DFT."
    )

    doc.add_heading("2.2 Quantum Transport Theory", level=2)

    p = doc.add_paragraph(
        "Transport in mesoscopic systems is governed by the Landauer-Buttiker formalism, which relates conductance "
        "to transmission probabilities through a scattering region "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[14])  # Patel2021
    p.add_run(
        ". In the quantum Hall regime, the Hall conductance is quantized in units of e2/h, reflecting the "
        "topological nature of the underlying electronic states. The theoretical framework for understanding "
        "these phenomena in strongly correlated systems was developed by "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[13])  # Nakamura2019
    p.add_run(".")

    doc.add_page_break()

    # --- Chapter 3: Experimental Methods ---
    doc.add_heading("Chapter 3: Experimental Methods", level=1)

    doc.add_heading("3.1 Sample Preparation", level=2)

    p = doc.add_paragraph(
        "High-quality MoS2 single crystals were grown by the chemical vapor transport (CVT) method following "
        "established protocols "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[9])  # Johnson2018
    p.add_run(
        ". Atomically thin flakes were obtained by mechanical exfoliation onto Si/SiO2 substrates with a 285 nm "
        "oxide layer, optimized for optical contrast. The layer number was confirmed by atomic force microscopy (AFM) "
        "and Raman spectroscopy."
    )

    doc.add_heading("3.2 Characterization Techniques", level=2)

    p = doc.add_paragraph(
        "Synchrotron-based X-ray characterization was performed at beamline 12.3.2 of the Advanced Light Source, "
        "following the methodology described in "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[4])  # Evans2017
    p.add_run(
        ". Micro-focused X-ray diffraction provided spatially resolved structural information with a beam size "
        "of approximately 1 micrometer, enabling correlation of local crystal structure with transport properties."
    )

    p2 = doc.add_paragraph(
        "Nanoscale characterization was performed using scanning tunneling microscopy and spectroscopy "
        "(STM/STS) at 4.2 K in ultra-high vacuum conditions "
    )
    add_citation_field(p2, BIBLIOGRAPHY_ENTRIES[6])  # Garcia2020
    p2.add_run(
        ". The STM measurements followed the protocols established in "
    )
    add_citation_field(p2, BIBLIOGRAPHY_ENTRIES[7])  # Hoffman2019
    p2.add_run(
        ", with particular attention to tip preparation and calibration procedures."
    )

    doc.add_page_break()

    # --- Chapter 4: Results and Discussion ---
    doc.add_heading("Chapter 4: Results and Discussion", level=1)

    doc.add_heading("4.1 Band Structure Calculations", level=2)

    p = doc.add_paragraph(
        "Our DFT calculations reveal a direct bandgap of 1.83 eV for monolayer MoS2, in excellent agreement "
        "with experimental photoluminescence measurements. The spin-orbit splitting at the valence band K point "
        "is calculated to be 148 meV, consistent with the theoretical predictions of "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[3])  # Davidson2021
    p.add_run(
        ". The inclusion of van der Waals corrections modifies the interlayer distance by approximately 3%, "
        "which has a negligible effect on the electronic structure of isolated monolayers."
    )

    doc.add_heading("4.2 Energy Applications", level=2)

    p = doc.add_paragraph(
        "The potential of these 2D materials for energy harvesting and storage applications has been extensively "
        "explored in recent literature "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[10])  # Kim2023
    p.add_run(
        ". Our measurements of the thermoelectric power factor in MoS2/WSe2 heterostructures demonstrate "
        "values exceeding 500 microWatts per meter-Kelvin-squared at room temperature, suggesting promising "
        "applications in waste heat recovery."
    )

    doc.add_heading("4.3 Spin Glass Behavior", level=2)

    p = doc.add_paragraph(
        "At low temperatures below 2 K, we observe signatures of spin glass freezing in heavily doped samples, "
        "consistent with the theoretical framework of "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[5])  # Fischer2022
    p.add_run(
        ". The frequency-dependent AC susceptibility measurements show a characteristic Vogel-Fulcher "
        "divergence of the relaxation time, with a freezing temperature of 1.4 K."
    )

    doc.add_heading("4.4 Ultrafast Dynamics", level=2)

    p = doc.add_paragraph(
        "Time-resolved pump-probe spectroscopy measurements, performed using the techniques described in "
    )
    add_citation_field(p, BIBLIOGRAPHY_ENTRIES[12])  # Martinez2022
    p.add_run(
        ", reveal ultrafast carrier dynamics with a characteristic relaxation time of 350 femtoseconds "
        "for the hot carrier cooling process. The valley polarization lifetime is measured to be 4.2 picoseconds "
        "at 10 K, significantly longer than previous reports in exfoliated samples."
    )

    doc.add_page_break()

    # --- Chapter 5: Conclusions ---
    doc.add_heading("Chapter 5: Conclusions and Future Work", level=1)

    p = doc.add_paragraph(
        "In this dissertation, we have presented a comprehensive study of the quantum transport properties "
        "of novel two-dimensional materials. Our key findings include the observation of fractional quantum Hall "
        "states in high-mobility MoS2, the discovery of spin glass behavior in heavily doped TMD samples, "
        "and the demonstration of ultrafast valley dynamics with record polarization lifetimes. These results "
        "advance our understanding of quantum phenomena in reduced dimensions and open new pathways for "
        "technological applications in quantum computing and energy harvesting."
    )

    p2 = doc.add_paragraph(
        "Future work will focus on extending these measurements to heterostructures combining TMDs with "
        "topological insulators, where the interplay of topology and strong correlations is expected to "
        "produce exotic quantum states of matter."
    )

    # NOTE: No bibliography index is inserted here. That is the agent's task.

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


create_initial()
