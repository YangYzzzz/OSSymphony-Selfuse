"""
Reward Script: Generate bibliography/references index at the end of thesis document
Task ID: writer_mt_077
Domain: libreoffice_writer
Scoring:
  Component 1: Bibliography heading exists with Heading 1 style (0.2 pts)
  Component 2: At least 15 bibliography entries after the heading (0.3 pts)
  Component 3: Entries in correct format Author, X. (Year). Title. Publisher. (0.25 pts)
  Component 4: Entries sorted alphabetically by author last name (0.25 pts)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_077'


def persist_app_state(domain: str):
    """Best-effort save for any open LibreOffice instance."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Find the Bibliography heading (must be Heading 1 style)
    bib_heading_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'bibliography' in p.text.strip().lower() and 'heading' in p.style.name.lower():
            bib_heading_idx = i
            break

    # Component 1: Bibliography heading exists with Heading 1 style (0.2 points)
    try:
        if bib_heading_idx >= 0:
            heading_style = doc.paragraphs[bib_heading_idx].style.name
            heading_text = doc.paragraphs[bib_heading_idx].text.strip()
            if heading_style == 'Heading 1':
                print(f"PASS: Component 1 — Bibliography heading found at para [{bib_heading_idx}] "
                      f"with style '{heading_style}', text='{heading_text}' (0.2 pts)")
                total_score += 0.2
            else:
                print(f"FAIL: Component 1 — Found heading '{heading_text}' but style is "
                      f"'{heading_style}', expected 'Heading 1'")
        else:
            print("FAIL: Component 1 — No 'Bibliography' heading with Heading style found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # If no bibliography heading, remaining checks cannot pass
    if bib_heading_idx < 0:
        print(f"\nScore: {total_score}/1.0")
        print(f"REWARD: {total_score}")
        return total_score

    # Collect bibliography entries (non-empty paragraphs after the heading)
    bib_entries = []
    for p in doc.paragraphs[bib_heading_idx + 1:]:
        txt = p.text.strip()
        if txt:
            bib_entries.append(txt)

    # Component 2: At least 15 bibliography entries present (0.3 points)
    # Progressive: partial credit for having some entries
    try:
        entry_count = len(bib_entries)
        if entry_count >= 15:
            print(f"PASS: Component 2 — {entry_count} bibliography entries found (>= 15) (0.3 pts)")
            total_score += 0.3
        elif entry_count >= 10:
            partial = 0.2
            print(f"PARTIAL: Component 2 — {entry_count} entries found (>= 10 but < 15) ({partial} pts)")
            total_score += partial
        elif entry_count >= 5:
            partial = 0.1
            print(f"PARTIAL: Component 2 — {entry_count} entries found (>= 5 but < 10) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — Only {entry_count} bibliography entries found, expected >= 15")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Entries in correct format "Author, X. (Year). Title. Publisher." (0.25 points)
    # Check that entries match the expected format pattern
    try:
        # Pattern: LastName, Initials (YYYY). Title. Publisher.
        # Initials can be "R. M." or "W." (one or two initials)
        format_pattern = re.compile(
            r'^[A-Z][a-z]+,\s+(?:[A-Z]\.\s*){1,3}\(\d{4}\)\.\s+.+\.\s+.+\.$'
        )
        formatted_count = 0
        for entry in bib_entries:
            if format_pattern.match(entry):
                formatted_count += 1

        if entry_count > 0:
            format_ratio = formatted_count / entry_count
        else:
            format_ratio = 0.0

        if format_ratio >= 0.9:
            print(f"PASS: Component 3 — {formatted_count}/{entry_count} entries in correct format "
                  f"({format_ratio:.0%}) (0.25 pts)")
            total_score += 0.25
        elif format_ratio >= 0.5:
            partial = round(0.25 * format_ratio, 2)
            print(f"PARTIAL: Component 3 — {formatted_count}/{entry_count} entries formatted correctly "
                  f"({format_ratio:.0%}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Only {formatted_count}/{entry_count} entries in correct format")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Entries sorted alphabetically by author last name (0.25 points)
    try:
        if entry_count >= 2:
            # Extract author last names for comparison
            last_names = []
            for entry in bib_entries:
                # Extract first word (last name) up to comma
                match = re.match(r'^([A-Za-z]+)', entry)
                if match:
                    last_names.append(match.group(1).lower())
                else:
                    last_names.append('')

            is_sorted = all(last_names[i] <= last_names[i + 1] for i in range(len(last_names) - 1))
            if is_sorted:
                print(f"PASS: Component 4 — Bibliography entries sorted alphabetically by author "
                      f"last name (0.25 pts)")
                total_score += 0.25
            else:
                # Check how many are in order
                in_order = sum(1 for i in range(len(last_names) - 1)
                               if last_names[i] <= last_names[i + 1])
                ratio = in_order / (len(last_names) - 1) if len(last_names) > 1 else 0
                print(f"FAIL: Component 4 — Entries not fully sorted. "
                      f"{in_order}/{len(last_names)-1} consecutive pairs in order")
        elif entry_count == 1:
            # Single entry is trivially sorted
            print(f"PASS: Component 4 — Single entry, trivially sorted (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 4 — No entries to check sort order")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Main entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
