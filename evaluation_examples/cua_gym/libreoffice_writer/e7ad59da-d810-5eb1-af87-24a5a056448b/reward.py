"""
Reward Script: Insert Cloud Storage adoption rate data from spreadsheet into Writer document
Task ID: osworld_multi_apps_calc_to_writer_013
Domain: libreoffice_writer (multi-app: calc + writer)
Scoring:
  - Component 1: Cloud Storage table exists in document (0.4 pts)
      A table with Cloud Storage data must be present — header row + at least 5 data rows
      where all data rows have Category='Cloud Storage'. This FAILS on initial_env (no such table).
  - Component 2: Data accuracy — all 5 Cloud Storage tools with correct adoption rates (0.4 pts)
      Checks exact tool names and adoption rate values match the spreadsheet source data.
      This FAILS on initial_env (no Cloud Storage table present).
  - Component 3: Placeholder text removed (0.2 pts)
      The placeholder '[Insert Cloud Storage adoption rate data from tech_adoption.xlsx here]'
      must no longer be present in the document. This FAILS on initial_env (placeholder still there).
"""

import os
from docx import Document

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_013'

# Expected Cloud Storage data from tech_adoption.xlsx (ground truth from spreadsheet)
EXPECTED_CLOUD_STORAGE = [
    ('Google Drive', 'Cloud Storage'),
    ('Microsoft OneDrive', 'Cloud Storage'),
    ('Dropbox', 'Cloud Storage'),
    ('Amazon S3', 'Cloud Storage'),
    ('Box', 'Cloud Storage'),
]

# Expected adoption rates (as they appear in the spreadsheet: 0.78, 0.65, etc.)
# Accept either decimal form (0.78), percent string (78%), or float string
EXPECTED_RATES = {
    'Google Drive': 0.78,
    'Microsoft OneDrive': 0.65,
    'Dropbox': 0.42,
    'Amazon S3': 0.53,
    'Box': 0.31,
}

PLACEHOLDER_TEXT = '[Insert Cloud Storage adoption rate data from tech_adoption.xlsx here]'


def parse_adoption_rate(value_str):
    """
    Parse adoption rate from various formats:
    - '78%' -> 0.78
    - '0.78' -> 0.78
    - 78 (int) -> 0.78 (if > 1, assume percentage)
    - 0.78 (float) -> 0.78
    """
    if value_str is None:
        return None
    s = str(value_str).strip()
    if s.endswith('%'):
        try:
            return float(s[:-1]) / 100.0
        except ValueError:
            return None
    try:
        val = float(s)
        # If value is > 1, assume it's a percentage (e.g. 78 means 78%)
        if val > 1.0:
            val = val / 100.0
        return val
    except ValueError:
        return None


def find_cloud_storage_table(doc):
    """
    Find a table in the document that contains Cloud Storage data.
    Returns the table if found, else None.
    A valid Cloud Storage table must have:
    - A header row with 'Technology' and 'Category' columns (or similar)
    - At least 1 data row with Category = 'Cloud Storage'
    """
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        # Check if any data row has 'Cloud Storage' category
        cloud_rows = 0
        for row in table.rows[1:]:  # skip header
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and 'Cloud Storage' in cells[1]:
                cloud_rows += 1
        if cloud_rows > 0:
            return table
    return None


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # -----------------------------------------------------------------------
    # Component 1: Cloud Storage table exists with correct structure (0.4 pts)
    # This FAILS on initial_env: no Cloud Storage table exists there.
    # -----------------------------------------------------------------------
    try:
        cloud_table = find_cloud_storage_table(doc)
        if cloud_table is not None:
            # Check it has at least 5 data rows (1 header + 5 data = 6 rows min)
            cloud_data_rows = 0
            for row in cloud_table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and 'Cloud Storage' in cells[1]:
                    cloud_data_rows += 1
            if cloud_data_rows >= 5:
                print(f"PASS: Component 1 — Cloud Storage table found with {cloud_data_rows} data rows (0.4 pts)")
                total_score += 0.4
            else:
                print(f"FAIL: Component 1 — Cloud Storage table found but only {cloud_data_rows} Cloud Storage rows (need 5+)")
                if cloud_data_rows > 0:
                    # Partial: at least a Cloud Storage table exists
                    partial = round(0.4 * cloud_data_rows / 5, 2)
                    print(f"  Partial credit: {partial} pts for {cloud_data_rows}/5 rows")
                    total_score += partial
        else:
            print("FAIL: Component 1 — No Cloud Storage table found in document")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # -----------------------------------------------------------------------
    # Component 2: Data accuracy — correct tool names and adoption rates (0.4 pts)
    # Each of the 5 tools is worth 0.08 pts (0.4 / 5).
    # This FAILS on initial_env: no Cloud Storage table to check.
    # -----------------------------------------------------------------------
    try:
        cloud_table = find_cloud_storage_table(doc)
        if cloud_table is not None:
            # Build a dict: technology_name -> adoption_rate from table
            found_tools = {}
            for row in cloud_table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 3 and 'Cloud Storage' in cells[1]:
                    tech_name = cells[0]
                    rate_raw = cells[2]
                    rate = parse_adoption_rate(rate_raw)
                    found_tools[tech_name] = rate

            correct_count = 0
            for tool, expected_rate in EXPECTED_RATES.items():
                if tool in found_tools:
                    actual_rate = found_tools[tool]
                    if actual_rate is not None and abs(actual_rate - expected_rate) < 0.01:
                        print(f"  PASS: {tool} rate={actual_rate:.2%} matches expected {expected_rate:.2%}")
                        correct_count += 1
                    else:
                        print(f"  FAIL: {tool} rate={actual_rate} expected {expected_rate:.2%}")
                else:
                    print(f"  FAIL: {tool} not found in Cloud Storage table")

            pts_earned = round(0.4 * correct_count / 5, 2)
            if correct_count == 5:
                print(f"PASS: Component 2 — All 5 tools with correct adoption rates ({pts_earned} pts)")
            else:
                print(f"FAIL: Component 2 — {correct_count}/5 tools with correct rates ({pts_earned} pts)")
            total_score += pts_earned
        else:
            print("FAIL: Component 2 — No Cloud Storage table found, cannot verify data accuracy")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # -----------------------------------------------------------------------
    # Component 3: Placeholder text removed (0.2 pts)
    # In initial_env, the placeholder '[Insert Cloud Storage adoption rate data...]' exists.
    # In golden_env, this placeholder should have been replaced by the table.
    # -----------------------------------------------------------------------
    try:
        placeholder_found = False
        for para in doc.paragraphs:
            if PLACEHOLDER_TEXT in para.text:
                placeholder_found = True
                break

        if not placeholder_found:
            print(f"PASS: Component 3 — Placeholder text removed from document (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Placeholder text still present: '{PLACEHOLDER_TEXT}'")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Default: test against canonical artifact path in golden/initial env
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
