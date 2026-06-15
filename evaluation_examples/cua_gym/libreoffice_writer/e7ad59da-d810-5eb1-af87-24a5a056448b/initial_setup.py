"""
Initial Setup: Technology review document with spreadsheet containing technology adoption data.
Task ID: osworld_multi_apps_calc_to_writer_013
Domain: libreoffice_writer (multi-app: Writer + Calc)

Creates:
  - /home/user/osworld_multi_apps_calc_to_writer_013.docx   (Writer review document)
  - /home/user/Desktop/tech_adoption.xlsx                    (spreadsheet with adoption data)
"""

import os
import shlex
import subprocess
import time

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_calc_to_writer_013'
DOC_OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
XLS_OUTPUT = f'{WORKDIR}/Desktop/tech_adoption.xlsx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_spreadsheet():
    """Create tech_adoption.xlsx on Desktop with multiple technology categories including Cloud Storage."""
    os.makedirs(f'{WORKDIR}/Desktop', exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Technology Adoption'

    # Style header row
    header_font = Font(name='Calibri', bold=True, size=12, color='FFFFFF')
    header_fill = PatternFill(start_color='FF2E4057', end_color='FF2E4057', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center')

    headers = ['Technology', 'Category', 'Adoption Rate', 'Year']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    # Column widths
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 22
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 10

    # Realistic technology adoption data
    # Categories: Cloud Storage, AI/ML Tools, DevOps Tools, Collaboration Tools, Security Tools
    data = [
        # Cloud Storage
        ('Google Drive', 'Cloud Storage', 0.78, 2024),
        ('Microsoft OneDrive', 'Cloud Storage', 0.65, 2024),
        ('Dropbox', 'Cloud Storage', 0.42, 2024),
        ('Amazon S3', 'Cloud Storage', 0.53, 2024),
        ('Box', 'Cloud Storage', 0.31, 2024),

        # AI/ML Tools
        ('GitHub Copilot', 'AI/ML Tools', 0.47, 2024),
        ('ChatGPT API', 'AI/ML Tools', 0.61, 2024),
        ('TensorFlow', 'AI/ML Tools', 0.38, 2024),
        ('Azure ML', 'AI/ML Tools', 0.29, 2024),
        ('Hugging Face', 'AI/ML Tools', 0.34, 2024),

        # DevOps Tools
        ('Docker', 'DevOps Tools', 0.72, 2024),
        ('Kubernetes', 'DevOps Tools', 0.56, 2024),
        ('Jenkins', 'DevOps Tools', 0.44, 2024),
        ('GitHub Actions', 'DevOps Tools', 0.68, 2024),
        ('Terraform', 'DevOps Tools', 0.41, 2024),

        # Collaboration Tools
        ('Slack', 'Collaboration Tools', 0.69, 2024),
        ('Microsoft Teams', 'Collaboration Tools', 0.74, 2024),
        ('Zoom', 'Collaboration Tools', 0.82, 2024),
        ('Notion', 'Collaboration Tools', 0.45, 2024),
        ('Confluence', 'Collaboration Tools', 0.38, 2024),

        # Security Tools
        ('Okta', 'Security Tools', 0.43, 2024),
        ('CrowdStrike', 'Security Tools', 0.36, 2024),
        ('Splunk', 'Security Tools', 0.31, 2024),
        ('Vault by HashiCorp', 'Security Tools', 0.27, 2024),
        ('Duo Security', 'Security Tools', 0.39, 2024),
    ]

    for r, row_data in enumerate(data, 2):
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val)
            if c == 3:
                cell.number_format = '0.00%'
            cell.alignment = Alignment(horizontal='left' if c <= 2 else 'center')

    # Freeze header row
    ws.freeze_panes = 'A2'

    wb.save(XLS_OUTPUT)
    print(f'Spreadsheet created: {XLS_OUTPUT}')


def create_writer_document():
    """Create the technology review document WITHOUT a table in the Technology Adoption Rates section."""
    doc = Document()

    # Title
    title = doc.add_heading('Enterprise Technology Review 2024', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Author / Date line
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Prepared by: IT Strategy Team  |  Date: November 2024')
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    # Executive Summary section
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents a comprehensive analysis of enterprise technology adoption trends '
        'across key digital transformation domains. The findings are based on surveys conducted '
        'with 1,200 IT professionals from organizations of varying sizes across North America, '
        'Europe, and the Asia-Pacific region.'
    )
    doc.add_paragraph(
        'Overall adoption of cloud-first strategies has accelerated significantly, with '
        'collaboration tools and cloud storage solutions leading adoption rates. Security '
        'and compliance considerations remain primary concerns for enterprise decision-makers.'
    )

    doc.add_paragraph()  # spacing

    # Key Findings section
    doc.add_heading('Key Findings', level=1)

    findings = [
        'Cloud storage solutions have achieved widespread enterprise adoption, with leading '
        'providers securing strong positions across organizations of all sizes.',
        'AI and machine learning tools are experiencing rapid growth, particularly in software '
        'development and data analytics use cases.',
        'DevOps tooling continues to mature, with containerization and infrastructure-as-code '
        'becoming standard practice in modern engineering organizations.',
        'Collaboration platforms have maintained high adoption rates following the global shift '
        'to hybrid and remote work models.',
        'Security tool adoption shows steady growth, driven by increasing regulatory requirements '
        'and the evolving threat landscape.',
    ]

    for finding in findings:
        p = doc.add_paragraph(finding, style='List Bullet')

    doc.add_paragraph()  # spacing

    # Technology Adoption Rates section
    doc.add_heading('Technology Adoption Rates', level=1)
    doc.add_paragraph(
        'The following section summarizes technology adoption rates across major categories '
        'surveyed in 2024. Adoption rates reflect the percentage of organizations actively '
        'deploying each technology in production environments.'
    )

    # Cloud Storage subsection - NO TABLE (agent needs to insert it)
    doc.add_heading('Cloud Storage', level=2)
    doc.add_paragraph(
        'Cloud storage solutions have seen broad adoption across enterprises, driven by '
        'scalability, cost efficiency, and remote collaboration requirements. Key vendors '
        'include both consumer-originated platforms that have expanded into enterprise markets '
        'and purpose-built enterprise storage solutions.'
    )

    # Placeholder note (agent should replace this content with the actual table)
    placeholder = doc.add_paragraph()
    run = placeholder.add_run('[Insert Cloud Storage adoption rate data from tech_adoption.xlsx here]')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacing

    # AI/ML Tools subsection
    doc.add_heading('AI and Machine Learning Tools', level=2)
    doc.add_paragraph(
        'Adoption of AI-powered tools has accelerated dramatically in 2024, with developer '
        'productivity tools such as code assistants leading the charge. Enterprise AI platforms '
        'show steady but more measured adoption due to governance and compliance considerations.'
    )

    table_aiml = doc.add_table(rows=1, cols=4)
    table_aiml.style = 'Table Grid'
    hdr_cells = table_aiml.rows[0].cells
    for i, h in enumerate(['Technology', 'Category', 'Adoption Rate', 'Year']):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    aiml_data = [
        ('GitHub Copilot', 'AI/ML Tools', '47%', '2024'),
        ('ChatGPT API', 'AI/ML Tools', '61%', '2024'),
        ('TensorFlow', 'AI/ML Tools', '38%', '2024'),
        ('Azure ML', 'AI/ML Tools', '29%', '2024'),
        ('Hugging Face', 'AI/ML Tools', '34%', '2024'),
    ]
    for row_data in aiml_data:
        row_cells = table_aiml.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph()  # spacing

    # DevOps section
    doc.add_heading('DevOps and Infrastructure Tools', level=2)
    doc.add_paragraph(
        'DevOps tooling has matured significantly, with containerization technologies and '
        'CI/CD automation becoming near-universal in engineering organizations. Infrastructure '
        'as code continues to gain traction as cloud infrastructure complexity increases.'
    )

    table_devops = doc.add_table(rows=1, cols=4)
    table_devops.style = 'Table Grid'
    hdr_cells = table_devops.rows[0].cells
    for i, h in enumerate(['Technology', 'Category', 'Adoption Rate', 'Year']):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    devops_data = [
        ('Docker', 'DevOps Tools', '72%', '2024'),
        ('Kubernetes', 'DevOps Tools', '56%', '2024'),
        ('Jenkins', 'DevOps Tools', '44%', '2024'),
        ('GitHub Actions', 'DevOps Tools', '68%', '2024'),
        ('Terraform', 'DevOps Tools', '41%', '2024'),
    ]
    for row_data in devops_data:
        row_cells = table_devops.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph()  # spacing

    # Collaboration section
    doc.add_heading('Collaboration and Communication Platforms', level=2)
    doc.add_paragraph(
        'Collaboration platforms have maintained consistently high adoption rates, reinforced '
        'by hybrid work policies and the need for seamless communication across distributed teams.'
    )

    doc.add_paragraph()  # spacing

    # Recommendations section
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph(
        'Based on the adoption trends identified in this report, the IT Strategy Team '
        'recommends the following actions for enterprise technology planning:'
    )

    recommendations = [
        'Prioritize cloud storage standardization to reduce fragmentation and improve '
        'data governance across the organization.',
        'Invest in AI tool training and governance frameworks to enable safe and '
        'productive adoption of generative AI tools.',
        'Accelerate DevOps maturity programs to close the gap between leading and '
        'lagging teams within the organization.',
        'Evaluate consolidation opportunities in the collaboration platform portfolio '
        'to reduce licensing costs and integration complexity.',
        'Strengthen security tooling investments to keep pace with evolving threats '
        'and compliance requirements.',
    ]

    for rec in recommendations:
        doc.add_paragraph(rec, style='List Number')

    doc.add_paragraph()  # spacing

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The 2024 technology adoption landscape reflects an organization in active '
        'transformation. Cloud and collaboration tools have achieved mainstream status, '
        'while AI and security tools represent the next wave of enterprise investment. '
        'Continued monitoring and strategic planning will be essential to ensure the '
        'organization remains competitive and secure.'
    )

    doc.save(DOC_OUTPUT)
    print(f'Writer document created: {DOC_OUTPUT}')


def main():
    # 1. Create the spreadsheet on Desktop
    create_spreadsheet()

    # 2. Create the Writer document
    create_writer_document()

    # 3. GUI-ready startup: open Writer document first, then Calc spreadsheet
    launch_gui(f'libreoffice --writer "{DOC_OUTPUT}"', delay_sec=3.0)
    launch_gui(f'libreoffice --calc "{XLS_OUTPUT}"', delay_sec=2.0)

    print('GUI_READY: Launched LibreOffice Writer and LibreOffice Calc with DISPLAY=:0')


main()
