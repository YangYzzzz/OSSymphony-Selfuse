"""
Initial Setup: Create a document with 8 bulleted takeaway items
Task ID: wrpara_039
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor

WORKDIR = '/home/user'
TASK_ID = 'wrpara_039'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add heading
    heading = doc.add_heading('Key Takeaways', level=1)

    # 8 bulleted items - realistic business/technology takeaways
    bullet_items = [
        "Cloud migration projects require at least six months of planning before the first workload is moved. Rushing the assessment phase leads to unexpected downtime and cost overruns that exceed initial budget estimates by 30 to 50 percent.",
        "Cross-functional teams consistently outperform siloed departments when tackling product launches. The blend of engineering, marketing, and customer success perspectives reduces time-to-market by an average of three weeks.",
        "Data governance policies must be established before scaling any analytics initiative. Without clear ownership and quality standards, dashboards become unreliable and decision-makers lose trust in the numbers.",
        "Remote onboarding effectiveness improves dramatically when new hires are paired with a dedicated mentor for the first 90 days. Retention rates among mentored employees are 25 percent higher than those onboarded through self-guided materials alone.",
        "API versioning should follow semantic conventions from day one, even for internal services. Retrofitting version management after multiple consumers depend on undocumented contracts creates cascading integration failures.",
        "Sustainability reporting is shifting from a compliance exercise to a strategic differentiator. Companies that publish transparent carbon reduction roadmaps attract institutional investors at a measurably higher rate.",
        "Automated regression testing catches roughly 80 percent of defects introduced during sprint development. The remaining 20 percent typically involve edge cases that require exploratory testing by experienced QA engineers.",
        "Customer feedback loops should be embedded directly into the product development cycle rather than handled as a separate quarterly review. Real-time sentiment analysis tools now make continuous feedback integration both feasible and cost-effective.",
    ]

    for item in bullet_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
