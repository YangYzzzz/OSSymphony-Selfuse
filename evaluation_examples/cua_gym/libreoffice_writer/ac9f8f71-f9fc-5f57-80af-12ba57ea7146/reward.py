"""
Reward Script: Enable header with bottom border in Default Page Style
Task ID: writer_fs_085
Domain: libreoffice_writer
Scoring:
  Component 1 (0.35): Header paragraph contains non-empty text
  Component 2 (0.35): Header paragraph has bottom border with single style and gray #808080 color
  Component 3 (0.30): Bottom border has correct width (0.25 pt) and spacing (~0.2 cm)
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_085'


def persist_app_state(domain):
    """Save any unsaved edits in LibreOffice before verification."""
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        import time
        time.sleep(1.0)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Precondition: document must have at least one section
    if len(doc.sections) == 0:
        print("FAIL: No sections found in document")
        print("REWARD: 0.0")
        return 0.0

    section = doc.sections[0]
    header = section.header

    # Precondition: header part must exist
    if header is None or len(header.paragraphs) == 0:
        print("FAIL: No header paragraphs found")
        print("REWARD: 0.0")
        return 0.0

    header_para = header.paragraphs[0]

    # Component 1: Header paragraph contains non-empty text (0.35 points)
    # Initial env has empty header text; golden env has actual header text.
    try:
        header_text = header_para.text.strip()
        if len(header_text) > 0:
            print(f"PASS: Component 1 — Header has text: '{header_text}' (0.35 pts)")
            total_score += 0.35
        else:
            print(f"FAIL: Component 1 — Header text is empty")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: Header paragraph has bottom border with single style and gray color (0.35 points)
    # In initial env there is no pBdr element; in golden env there is w:bottom with val=single, color=808080
    try:
        pPr = header_para._element.find(qn('w:pPr'))
        border_found = False
        border_style_ok = False
        border_color_ok = False

        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                bottom = pBdr.find(qn('w:bottom'))
                if bottom is not None:
                    border_found = True
                    val = bottom.get(qn('w:val'))
                    color = bottom.get(qn('w:color'))

                    # Check style is single (solid)
                    if val is not None and val.lower() == 'single':
                        border_style_ok = True

                    # Check color is gray #808080 (case-insensitive)
                    if color is not None and color.lower() == '808080':
                        border_color_ok = True

                    print(f"  Bottom border: val={val}, color={color}")

        if border_found and border_style_ok and border_color_ok:
            print(f"PASS: Component 2 — Bottom border is single style with #808080 color (0.35 pts)")
            total_score += 0.35
        elif border_found:
            # Partial: border exists but style or color wrong
            partial = 0.0
            if border_style_ok:
                partial += 0.15
            if border_color_ok:
                partial += 0.15
            print(f"PARTIAL: Component 2 — Border found but style_ok={border_style_ok}, color_ok={border_color_ok} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No bottom border found on header paragraph")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Bottom border has correct width (0.25 pt = sz 2) and spacing (~0.2 cm ≈ 6 pt) (0.30 points)
    # Initial env has no border at all, so this will fail on initial.
    try:
        pPr = header_para._element.find(qn('w:pPr'))
        sz_ok = False
        space_ok = False

        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                bottom = pBdr.find(qn('w:bottom'))
                if bottom is not None:
                    sz = bottom.get(qn('w:sz'))
                    space = bottom.get(qn('w:space'))

                    # w:sz is in eighths of a point. 0.25 pt = 2 eighths.
                    # Accept sz of 2 (0.25 pt) or close values (1-4 for tolerance)
                    if sz is not None:
                        sz_val = int(sz)
                        if sz_val == 2:
                            sz_ok = True
                        elif 1 <= sz_val <= 4:
                            # Thin border, close enough
                            sz_ok = True
                            print(f"  Note: sz={sz_val} (expected 2 for 0.25pt, accepting as close)")

                    # w:space is in points. 0.2 cm ≈ 5.67 pt. Accept 4-8 range.
                    if space is not None:
                        space_val = int(space)
                        if 4 <= space_val <= 8:
                            space_ok = True

                    print(f"  Border details: sz={sz}, space={space}")

        if sz_ok and space_ok:
            print(f"PASS: Component 3 — Border width and spacing correct (0.30 pts)")
            total_score += 0.30
        elif sz_ok or space_ok:
            partial = 0.15
            print(f"PARTIAL: Component 3 — sz_ok={sz_ok}, space_ok={space_ok} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — Border width/spacing not found or incorrect")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
