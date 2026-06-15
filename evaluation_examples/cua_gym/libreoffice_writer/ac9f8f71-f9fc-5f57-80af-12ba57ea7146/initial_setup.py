"""
Initial Setup: Enable header with bottom border in Writer document
Task ID: writer_fs_085
Domain: libreoffice_writer

Creates a 6-page Writer document with realistic business content.
No header is enabled -- the agent must enable it and add the border.
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_fs_085'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style.font.size = Pt(12)

    # -- Page 1: Title and Introduction --
    title = doc.add_heading('Quarterly Performance Review - Q1 2025', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('')  # spacer

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Prepared by: Human Resources Department')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = date_para.add_run('Date: March 28, 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph('')

    doc.add_paragraph(
        'This report summarizes the performance metrics and key achievements '
        'of the Engineering, Marketing, and Sales departments during the first '
        'quarter of fiscal year 2025. The analysis covers revenue targets, '
        'project milestones, employee engagement scores, and strategic initiatives '
        'undertaken during January through March.'
    )

    doc.add_paragraph(
        'The overall performance trajectory indicates strong growth across most '
        'divisions, with particular success in the cloud infrastructure migration '
        'project and the North American market expansion campaign. Areas requiring '
        'attention include staffing levels in the QA team and the delayed rollout '
        'of the mobile application redesign.'
    )

    # -- Page break to page 2 --
    doc.add_page_break()

    # -- Page 2: Engineering Department --
    doc.add_heading('Engineering Department', level=2)

    doc.add_paragraph(
        'The Engineering department delivered 14 out of 17 planned features during '
        'Q1, achieving an 82% completion rate. Lead engineer Sarah Chen spearheaded '
        'the cloud migration initiative, successfully transitioning 340 microservices '
        'from on-premise infrastructure to AWS. The migration reduced operational '
        'costs by $127,000 per month and improved average response times by 23%.'
    )

    doc.add_paragraph(
        'Notable technical achievements include the implementation of a new CI/CD '
        'pipeline that reduced deployment times from 45 minutes to under 8 minutes. '
        'The team also resolved 89 critical bugs from the backlog, bringing the '
        'defect density down to 0.7 per thousand lines of code.'
    )

    doc.add_heading('Key Metrics', level=3)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Metric', 'Target', 'Actual']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    data = [
        ['Features Delivered', '17', '14'],
        ['Bug Resolution Rate', '85%', '91%'],
        ['Code Coverage', '80%', '83.4%'],
        ['Deployment Frequency', '2x/week', '3x/week'],
        ['System Uptime', '99.9%', '99.97%'],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # -- Page break to page 3 --
    doc.add_page_break()

    # -- Page 3: Marketing Department --
    doc.add_heading('Marketing Department', level=2)

    doc.add_paragraph(
        'Under the leadership of Marcus Johnson, the Marketing team executed 23 '
        'campaigns across digital and traditional channels. The North American '
        'expansion campaign generated 4,200 qualified leads, surpassing the target '
        'of 3,500 by 20%. Social media engagement increased by 34% quarter-over-quarter, '
        'driven primarily by the LinkedIn thought leadership series and the redesigned '
        'company blog.'
    )

    doc.add_paragraph(
        'The content marketing strategy yielded exceptional results, with organic '
        'search traffic increasing by 47%. The team published 36 blog posts, 8 '
        'whitepapers, and produced 12 webinars that collectively attracted over '
        '15,000 registrations. Customer acquisition cost decreased from $342 to '
        '$287, representing a 16% improvement.'
    )

    doc.add_paragraph(
        'Brand awareness metrics showed positive trends across all tracked channels. '
        'The Net Promoter Score improved from 42 to 51, indicating stronger customer '
        'satisfaction and loyalty. The partnership with TechInsight Media for sponsored '
        'content contributed to a 28% increase in enterprise-level inquiries.'
    )

    # -- Page break to page 4 --
    doc.add_page_break()

    # -- Page 4: Sales Department --
    doc.add_heading('Sales Department', level=2)

    doc.add_paragraph(
        'The Sales department achieved $4.2 million in revenue against a target of '
        '$3.8 million, representing a 110% attainment rate. Regional director Elena '
        'Rodriguez led the enterprise sales team to close 7 deals exceeding $200,000 '
        'each, including the landmark partnership with GlobalTech Industries valued '
        'at $780,000 annually.'
    )

    doc.add_heading('Revenue Breakdown by Region', level=3)

    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    headers2 = ['Region', 'Target', 'Actual', 'Attainment']
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    sales_data = [
        ['North America', '$1,500,000', '$1,720,000', '115%'],
        ['Europe', '$1,200,000', '$1,180,000', '98%'],
        ['Asia-Pacific', '$800,000', '$920,000', '115%'],
        ['Latin America', '$300,000', '$380,000', '127%'],
    ]
    for r, row_data in enumerate(sales_data, 1):
        for c, val in enumerate(row_data):
            table2.cell(r, c).text = val

    doc.add_paragraph(
        'Pipeline development remains healthy with $6.8 million in qualified '
        'opportunities for Q2. The sales enablement program trained 12 new account '
        'executives, and the revised commission structure has improved team retention '
        'rates to 94%.'
    )

    # -- Page break to page 5 --
    doc.add_page_break()

    # -- Page 5: Employee Engagement & HR --
    doc.add_heading('Employee Engagement & Human Resources', level=2)

    doc.add_paragraph(
        'The quarterly employee engagement survey, conducted in March, received an '
        '87% response rate across all departments. Overall engagement score rose from '
        '7.2 to 7.8 out of 10, reflecting the positive impact of recent workplace '
        'initiatives including flexible work arrangements, the new mentorship program, '
        'and upgraded break room facilities.'
    )

    doc.add_paragraph(
        'Hiring activity during Q1 resulted in 28 new employees across 5 departments. '
        'Time-to-fill metrics improved from 38 days to 29 days, attributed to the '
        'partnership with three new recruitment agencies and the enhanced employee '
        'referral bonus program. The voluntary attrition rate decreased from 4.1% to '
        '2.8%, well below the industry average of 5.3%.'
    )

    doc.add_paragraph(
        'Training and development programs saw record participation, with 156 employees '
        'completing at least one professional development course. The leadership '
        'development track graduated its second cohort of 15 emerging leaders, three '
        'of whom have since been promoted to management positions.'
    )

    # -- Page break to page 6 --
    doc.add_page_break()

    # -- Page 6: Strategic Outlook --
    doc.add_heading('Strategic Outlook for Q2 2025', level=2)

    doc.add_paragraph(
        'Looking ahead to Q2, the organization will focus on three strategic pillars: '
        'completing the cloud infrastructure migration, launching the redesigned mobile '
        'application, and expanding into the Southeast Asian market. The board has '
        'approved a $1.2 million budget increase for the technology modernization program.'
    )

    doc.add_paragraph(
        'Key priorities include establishing a regional office in Singapore, hiring '
        '15 additional engineers for the platform team, and launching the customer '
        'success program to reduce churn by an estimated 20%. The marketing team will '
        'debut the new brand identity at the TechForward Conference in June.'
    )

    doc.add_paragraph(
        'Risk factors to monitor include potential supply chain disruptions affecting '
        'hardware procurement, competitive pressure from emerging SaaS providers, and '
        'regulatory changes in European data privacy requirements. Mitigation strategies '
        'for each risk area are documented in the appendix to this report.'
    )

    doc.add_paragraph('')
    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = closing.add_run('Report compiled by the Office of the Chief Operating Officer')
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Ensure NO header is set (python-docx default is no header, but be explicit)
    for section in doc.sections:
        section.header.is_linked_to_previous = True
        # Clear any default header paragraphs
        for para in section.header.paragraphs:
            para.text = ''

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
