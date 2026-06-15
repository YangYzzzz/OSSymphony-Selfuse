"""
Initial Setup: Create a document with acronyms in regular text (no character styles, no small caps)
Task ID: writer_bs_086
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_086'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading('International Organizations and Their Impact', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Paragraph 1 - NASA
    doc.add_paragraph(
        'The National Aeronautics and Space Administration, commonly known as NASA, '
        'has been at the forefront of space exploration since its founding in 1958. '
        'NASA has launched numerous missions including the Apollo program, the Space '
        'Shuttle program, and the Mars Rover expeditions. Today, NASA continues to '
        'push the boundaries of human knowledge through the Artemis program and '
        'partnerships with private aerospace companies.'
    )

    # Paragraph 2 - UNESCO
    doc.add_paragraph(
        'The United Nations Educational, Scientific and Cultural Organization, or UNESCO, '
        'plays a vital role in promoting international cooperation in education, science, '
        'and culture. UNESCO designates World Heritage Sites across the globe, protecting '
        'locations of outstanding cultural or natural importance. Since its establishment '
        'in 1945, UNESCO has been instrumental in promoting literacy and preserving '
        'cultural heritage in over 190 member states.'
    )

    # Paragraph 3 - NATO
    doc.add_paragraph(
        'The North Atlantic Treaty Organization, known as NATO, is an intergovernmental '
        'military alliance established in 1949. NATO was created to provide collective '
        'defense against potential threats during the Cold War era. Over the decades, '
        'NATO has expanded its membership and adapted its mission to address modern '
        'security challenges including cybersecurity and counterterrorism. NATO currently '
        'comprises 31 member nations committed to mutual defense.'
    )

    # Paragraph 4 - WHO
    doc.add_paragraph(
        'The World Health Organization, abbreviated as WHO, is a specialized agency of '
        'the United Nations responsible for international public health. WHO coordinates '
        'global health responses, sets health standards, and provides technical assistance '
        'to countries in need. During the COVID-19 pandemic, WHO played a central role in '
        'coordinating the international response and facilitating vaccine distribution. '
        'WHO also leads efforts to eradicate diseases such as malaria and tuberculosis.'
    )

    # Paragraph 5 - EU and cross-references
    doc.add_paragraph(
        'The European Union, referred to as the EU, is a political and economic union of '
        '27 member states located primarily in Europe. The EU has established a single '
        'market that allows free movement of goods, services, capital, and people. Many '
        'EU member states are also part of NATO, creating overlapping security frameworks. '
        'The EU works closely with WHO on public health initiatives and collaborates with '
        'UNESCO on cultural preservation projects. Additionally, the EU has partnered with '
        'NASA on several space research programs.'
    )

    # Concluding paragraph
    doc.add_paragraph(
        'These international organizations - NASA, UNESCO, NATO, WHO, and the EU - each '
        'serve distinct yet complementary roles in shaping global policy and advancing '
        'human progress. Their collaborative efforts demonstrate the importance of '
        'multilateral cooperation in addressing the complex challenges facing our world today.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
