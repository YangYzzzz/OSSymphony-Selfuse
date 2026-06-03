"""
Reward Script: Create 'Acronym' character style with small caps and apply to all acronyms
Task ID: writer_bs_086
Domain: libreoffice_writer
Scoring:
  Component 1 (0.3): 'Acronym' character style exists with smallCaps
  Component 2 (0.5): All acronym occurrences have the 'Acronym' char style applied
  Component 3 (0.2): Acronym-styled runs effectively have smallCaps formatting
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_086'

# The five acronyms that must have the Acronym style applied
ACRONYMS = ['NASA', 'UNESCO', 'NATO', 'WHO', 'EU']


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Missing dependency: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # --- Precondition gate: document text is intact ---
    expected_snippets = [
        "International Organizations and Their Impact",
        "National Aeronautics and Space Administration",
        "North Atlantic Treaty Organization",
    ]
    full_text = " ".join(para.text for para in doc.paragraphs)
    for snippet in expected_snippets:
        if snippet not in full_text:
            print(f"PRECONDITION FAIL: Missing text '{snippet}' — document corrupted")
            print("REWARD: 0.0")
            return 0.0

    # ---------------------------------------------------------------
    # Component 1: 'Acronym' character style exists with smallCaps (0.3 points)
    # This FAILS on initial_env (no 'Acronym' style) and PASSES on golden_env.
    # ---------------------------------------------------------------
    acronym_style_has_smallcaps = False
    try:
        for style in doc.styles:
            if style.name == 'Acronym':
                if style.type is not None and style.type.name == 'CHARACTER':
                    # Check for smallCaps in the style's run properties
                    rpr = style.element.find(qn('w:rPr'))
                    if rpr is not None:
                        sc_el = rpr.find(qn('w:smallCaps'))
                        if sc_el is not None:
                            val = sc_el.get(qn('w:val'))
                            # smallCaps present and not explicitly disabled
                            if val is None or val.lower() in ('true', '1', 'on'):
                                acronym_style_has_smallcaps = True
                break

        if acronym_style_has_smallcaps:
            print("PASS: Component 1 — 'Acronym' character style exists with smallCaps (0.3 pts)")
            total_score += 0.3
        else:
            print("FAIL: Component 1 — 'Acronym' character style with smallCaps not found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ---------------------------------------------------------------
    # Component 2: All acronym occurrences have the 'Acronym' char style (0.5 points)
    # This FAILS on initial_env (all text in single unstyled runs) and PASSES on golden_env.
    # ---------------------------------------------------------------
    try:
        # Count expected acronym occurrences in the document text
        expected_counts = {}
        for para in doc.paragraphs:
            text = para.text
            for acr in ACRONYMS:
                if acr == 'EU':
                    pattern = r'\bEU\b'
                else:
                    pattern = r'\b' + re.escape(acr) + r'\b'
                matches = re.findall(pattern, text)
                expected_counts[acr] = expected_counts.get(acr, 0) + len(matches)

        total_expected = sum(expected_counts.values())
        print(f"  Expected acronym occurrences: {expected_counts} (total: {total_expected})")

        # Count how many acronym runs have the 'Acronym' character style
        styled_count = 0

        for para in doc.paragraphs:
            for run in para.runs:
                run_text = run.text.strip()
                if run_text in ACRONYMS:
                    rpr = run._element.find(qn('w:rPr'))
                    if rpr is not None:
                        rstyle_el = rpr.find(qn('w:rStyle'))
                        if rstyle_el is not None:
                            style_val = rstyle_el.get(qn('w:val'))
                            if style_val == 'Acronym':
                                styled_count += 1

        print(f"  Acronym runs with 'Acronym' style: {styled_count}/{total_expected}")

        if total_expected > 0 and styled_count >= total_expected:
            print(f"PASS: Component 2 — All {styled_count}/{total_expected} acronym occurrences styled (0.5 pts)")
            total_score += 0.5
        elif total_expected > 0 and styled_count > 0:
            ratio = styled_count / total_expected
            partial = round(0.5 * ratio, 2)
            print(f"PARTIAL: Component 2 — {styled_count}/{total_expected} acronyms styled ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 — No acronym occurrences have 'Acronym' character style")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ---------------------------------------------------------------
    # Component 3: Acronym-styled runs have effective smallCaps (0.2 points)
    # Verifies that runs with 'Acronym' style actually render as small caps,
    # either via the style definition or via direct formatting on the run.
    # FAILS on initial_env (no runs with smallCaps), PASSES on golden_env.
    # ---------------------------------------------------------------
    try:
        acronym_runs_with_smallcaps = 0
        acronym_runs_total = 0

        for para in doc.paragraphs:
            for run in para.runs:
                run_text = run.text.strip()
                if run_text in ACRONYMS:
                    # Check if this run has Acronym style applied
                    rpr = run._element.find(qn('w:rPr'))
                    has_acronym_style = False
                    has_direct_smallcaps = False

                    if rpr is not None:
                        rstyle_el = rpr.find(qn('w:rStyle'))
                        if rstyle_el is not None and rstyle_el.get(qn('w:val')) == 'Acronym':
                            has_acronym_style = True

                        sc_el = rpr.find(qn('w:smallCaps'))
                        if sc_el is not None:
                            val = sc_el.get(qn('w:val'))
                            if val is None or val.lower() in ('true', '1', 'on'):
                                has_direct_smallcaps = True

                    # SmallCaps can come from either the style or direct formatting
                    if has_acronym_style or has_direct_smallcaps:
                        acronym_runs_total += 1
                        # If the style has smallCaps (from Component 1 check) or direct smallCaps
                        if (has_acronym_style and acronym_style_has_smallcaps) or has_direct_smallcaps:
                            acronym_runs_with_smallcaps += 1

        print(f"  Acronym runs with effective smallCaps: {acronym_runs_with_smallcaps}/{acronym_runs_total}")

        if acronym_runs_total > 0 and acronym_runs_with_smallcaps == acronym_runs_total:
            print(f"PASS: Component 3 — All acronym runs have effective smallCaps formatting (0.2 pts)")
            total_score += 0.2
        elif acronym_runs_with_smallcaps > 0:
            ratio = acronym_runs_with_smallcaps / max(acronym_runs_total, 1)
            partial = round(0.2 * ratio, 2)
            print(f"PARTIAL: Component 3 — {acronym_runs_with_smallcaps}/{acronym_runs_total} have smallCaps ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 — No acronym runs have effective smallCaps formatting")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook for LibreOffice Writer
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    persist_app_state()
    verify_task(file_path)
