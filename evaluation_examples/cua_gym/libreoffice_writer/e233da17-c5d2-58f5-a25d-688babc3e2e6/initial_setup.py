"""
Initial Setup: Insert a Table of Tables for a research paper
Task ID: writer_mt_057
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_057'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_field_code(paragraph, instr_text):
    """Add a field code to a paragraph (e.g., SEQ Table, PAGE, TOC)."""
    # Begin field char
    r_begin = paragraph._element.makeelement(qn('w:r'), {})
    fld_begin = r_begin.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r_begin.append(fld_begin)
    paragraph._element.append(r_begin)

    # Instruction text
    r_instr = paragraph._element.makeelement(qn('w:r'), {})
    instr_el = r_instr.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr_el.text = instr_text
    r_instr.append(instr_el)
    paragraph._element.append(r_instr)

    # Separate field char
    r_sep = paragraph._element.makeelement(qn('w:r'), {})
    fld_sep = r_sep.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'})
    r_sep.append(fld_sep)
    paragraph._element.append(r_sep)

    # Result placeholder
    r_result = paragraph._element.makeelement(qn('w:r'), {})
    t_result = r_result.makeelement(qn('w:t'), {})
    t_result.text = ''
    r_result.append(t_result)
    paragraph._element.append(r_result)

    # End field char
    r_end = paragraph._element.makeelement(qn('w:r'), {})
    fld_end = r_end.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r_end.append(fld_end)
    paragraph._element.append(r_end)


def add_caption_paragraph(doc, table_num, caption_text):
    """Add a proper caption paragraph with SEQ field: 'Table N: Caption Text'."""
    para = doc.add_paragraph()
    para.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else doc.styles['Normal']
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add 'Table ' text
    run_prefix = para.add_run('Table ')
    run_prefix.font.size = Pt(10)
    run_prefix.italic = True

    # Add SEQ Table field
    add_field_code(para, ' SEQ Table \\* ARABIC ')

    # Set the field result to the table number
    # Find the result run (second to last run element)
    runs = para._element.findall(qn('w:r'))
    # The result run is between separate and end
    for i, r in enumerate(runs):
        fld_chars = r.findall(qn('w:fldChar'))
        for fc in fld_chars:
            if fc.get(qn('w:fldCharType')) == 'separate':
                # Next run is the result
                result_run = runs[i + 1]
                t_el = result_run.find(qn('w:t'))
                if t_el is not None:
                    t_el.text = str(table_num)
                break

    # Add ': Caption Text'
    run_caption = para.add_run(f': {caption_text}')
    run_caption.font.size = Pt(10)
    run_caption.italic = True

    return para


def add_table_with_data(doc, headers, data):
    """Add a table with headers and data rows."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'

    # Headers
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # Data
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            table.cell(i + 1, j).text = str(val)

    return table


def create_initial():
    doc = Document()

    # ==================== Title Page ====================
    doc.add_heading('Impact of Remote Work on Employee Productivity\nand Well-Being: A Comprehensive Study', level=0)
    doc.add_paragraph('')
    authors = doc.add_paragraph()
    authors.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors.add_run('Dr. Elena Rodriguez, Ph.D.').bold = True
    authors.add_run('\nDepartment of Organizational Psychology\nStanford University')
    doc.add_paragraph('')
    date_para = doc.add_paragraph()
    date_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.add_run('March 2025')

    doc.add_page_break()

    # ==================== Abstract ====================
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This study investigates the multifaceted impact of remote work arrangements on employee '
        'productivity and psychological well-being across diverse industry sectors. Using a mixed-methods '
        'approach combining quantitative surveys (n=1,247) with semi-structured interviews (n=85), we '
        'examined how work-from-home policies implemented during and after the COVID-19 pandemic affected '
        'key performance indicators, job satisfaction, and mental health outcomes. Our findings reveal a '
        'nuanced picture: while overall productivity increased by 13.2% among remote workers, significant '
        'variations emerged across demographic groups, job types, and organizational cultures. The study '
        'contributes to the growing literature on flexible work arrangements by identifying critical '
        'moderating factors including managerial support, home office infrastructure, and social isolation.'
    )

    doc.add_page_break()

    # ==================== Introduction ====================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The rapid transition to remote work during the global pandemic fundamentally transformed '
        'workplace dynamics across industries. Organizations that previously relied on traditional '
        'office-based models were forced to adopt distributed work arrangements virtually overnight. '
        'This unprecedented shift created a natural experiment of enormous scale, offering researchers '
        'a unique opportunity to examine the effects of remote work under conditions that would have '
        'been impossible to replicate in controlled settings.'
    )
    doc.add_paragraph(
        'Prior research on telecommuting, while extensive, has largely focused on voluntary arrangements '
        'in knowledge-work contexts (Bailey & Kurland, 2002; Gajendran & Harrison, 2007). The pandemic-era '
        'shift differed fundamentally: it was involuntary, universal within organizations, and accompanied '
        'by unprecedented external stressors including health anxiety, school closures, and social isolation. '
        'Understanding how these conditions shaped the remote work experience is crucial for developing '
        'evidence-based policies as organizations navigate post-pandemic hybrid models.'
    )
    doc.add_paragraph(
        'This paper addresses three primary research questions: (1) How did the shift to remote work '
        'affect measurable productivity outcomes? (2) What demographic and organizational factors moderated '
        'these effects? (3) How did prolonged remote work impact employee psychological well-being? We '
        'examine these questions through comprehensive survey data and in-depth interviews, contributing '
        'both quantitative evidence and rich qualitative insights to the ongoing debate.'
    )

    doc.add_page_break()

    # ==================== Methods ====================
    doc.add_heading('2. Methods', level=1)

    doc.add_heading('2.1 Participants', level=2)
    doc.add_paragraph(
        'A total of 1,247 full-time employees participated in the quantitative survey phase. Participants '
        'were recruited from 23 organizations spanning technology, finance, healthcare administration, '
        'education, and professional services sectors. Inclusion criteria required participants to have '
        'worked both in-office and remotely for minimum periods of six months each during the study window '
        '(January 2022 to December 2024).'
    )

    # TABLE 1: Demographics
    add_table_with_data(doc,
        ['Characteristic', 'Category', 'n', 'Percentage'],
        [
            ['Age Group', '18-29', '287', '23.0%'],
            ['', '30-39', '412', '33.0%'],
            ['', '40-49', '318', '25.5%'],
            ['', '50+', '230', '18.4%'],
            ['Gender', 'Female', '648', '52.0%'],
            ['', 'Male', '571', '45.8%'],
            ['', 'Non-binary', '28', '2.2%'],
            ['Education', 'Bachelor\'s', '489', '39.2%'],
            ['', 'Master\'s', '524', '42.0%'],
            ['', 'Doctorate', '234', '18.8%'],
            ['Industry', 'Technology', '378', '30.3%'],
            ['', 'Finance', '267', '21.4%'],
            ['', 'Healthcare Admin', '198', '15.9%'],
            ['', 'Education', '221', '17.7%'],
            ['', 'Professional Services', '183', '14.7%'],
        ]
    )
    add_caption_paragraph(doc, 1, 'Demographics')
    doc.add_paragraph('')

    doc.add_heading('2.2 Survey Instrument', level=2)
    doc.add_paragraph(
        'The survey instrument comprised 72 items organized into four domains: productivity self-assessment '
        '(adapted from Endicott Work Productivity Scale), job satisfaction (Minnesota Satisfaction '
        'Questionnaire short form), psychological well-being (WHO-5 Well-Being Index), and remote work '
        'environment quality (custom scale developed through pilot testing, Cronbach\'s alpha = 0.89). '
        'All items used 5-point Likert scales unless otherwise specified.'
    )

    doc.add_page_break()

    # ==================== Results ====================
    doc.add_heading('3. Results', level=1)

    doc.add_heading('3.1 Survey Findings', level=2)
    doc.add_paragraph(
        'The primary survey yielded robust response rates across all participating organizations (mean '
        'response rate = 78.3%, range: 64.1% - 91.7%). After data cleaning and removal of incomplete '
        'responses, the final analytical sample comprised 1,247 participants.'
    )

    # TABLE 2: Survey Results
    add_table_with_data(doc,
        ['Domain', 'Measure', 'Remote Mean (SD)', 'Office Mean (SD)', 'p-value'],
        [
            ['Productivity', 'Task Completion Rate', '87.3 (12.1)', '79.6 (14.8)', '<0.001'],
            ['', 'Quality Rating', '4.21 (0.68)', '4.15 (0.72)', '0.042'],
            ['', 'Meeting Efficiency', '3.89 (0.91)', '3.45 (1.02)', '<0.001'],
            ['Satisfaction', 'Overall Job Satisfaction', '4.02 (0.85)', '3.78 (0.93)', '<0.001'],
            ['', 'Work-Life Balance', '3.94 (1.12)', '3.21 (1.08)', '<0.001'],
            ['', 'Career Growth Concern', '3.45 (1.18)', '2.87 (1.05)', '<0.001'],
            ['Well-Being', 'WHO-5 Score', '62.4 (18.7)', '58.9 (17.2)', '<0.001'],
            ['', 'Burnout Index', '2.87 (0.94)', '3.12 (0.88)', '0.003'],
            ['', 'Social Isolation', '3.56 (1.23)', '2.14 (0.98)', '<0.001'],
        ]
    )
    add_caption_paragraph(doc, 2, 'Survey Results')
    doc.add_paragraph('')

    doc.add_paragraph(
        'Results indicated statistically significant improvements in productivity metrics for remote workers '
        'compared to office-based periods. Task completion rates increased by 9.7% on average (t(1246) = '
        '8.34, p < .001, Cohen\'s d = 0.57), while quality ratings showed a modest but significant '
        'improvement (t(1246) = 2.03, p = .042, d = 0.09).'
    )

    # TABLE 3: Statistical Analysis
    add_table_with_data(doc,
        ['Variable', 'B', 'SE', 't', 'p', '95% CI'],
        [
            ['Intercept', '3.241', '0.187', '17.33', '<.001', '[2.874, 3.608]'],
            ['Remote Work (1=yes)', '0.482', '0.093', '5.18', '<.001', '[0.299, 0.665]'],
            ['Age', '-0.012', '0.004', '-3.00', '.003', '[-0.020, -0.004]'],
            ['Gender (1=female)', '0.156', '0.089', '1.75', '.080', '[-0.019, 0.331]'],
            ['Education (years)', '0.078', '0.021', '3.71', '<.001', '[0.037, 0.119]'],
            ['Manager Support', '0.341', '0.067', '5.09', '<.001', '[0.210, 0.472]'],
            ['Home Office Quality', '0.267', '0.058', '4.60', '<.001', '[0.153, 0.381]'],
            ['Social Isolation', '-0.198', '0.045', '-4.40', '<.001', '[-0.286, -0.110]'],
        ]
    )
    add_caption_paragraph(doc, 3, 'Statistical Analysis')
    doc.add_paragraph('')

    doc.add_page_break()

    doc.add_heading('3.2 Correlation Analysis', level=2)
    doc.add_paragraph(
        'Bivariate correlations among key study variables are presented below. Several noteworthy '
        'patterns emerged. Productivity showed moderate positive correlations with manager support '
        '(r = .42) and home office quality (r = .38), and a moderate negative correlation with social '
        'isolation (r = -.35).'
    )

    # TABLE 4: Correlation Matrix
    add_table_with_data(doc,
        ['', 'Productivity', 'Satisfaction', 'Well-Being', 'Mgr Support', 'Home Office', 'Isolation'],
        [
            ['Productivity', '1.00', '.54**', '.48**', '.42**', '.38**', '-.35**'],
            ['Satisfaction', '.54**', '1.00', '.61**', '.52**', '.31**', '-.29**'],
            ['Well-Being', '.48**', '.61**', '1.00', '.45**', '.27**', '-.51**'],
            ['Mgr Support', '.42**', '.52**', '.45**', '1.00', '.18**', '-.22**'],
            ['Home Office', '.38**', '.31**', '.27**', '.18**', '1.00', '-.15**'],
            ['Isolation', '-.35**', '-.29**', '-.51**', '-.22**', '-.15**', '1.00'],
        ]
    )
    add_caption_paragraph(doc, 4, 'Correlation Matrix')
    doc.add_paragraph('')

    doc.add_heading('3.3 Regression Analysis', level=2)
    doc.add_paragraph(
        'A hierarchical multiple regression analysis was conducted to examine the predictors of remote '
        'work productivity. The full model explained 47.3% of variance in productivity scores '
        '(F(7, 1239) = 158.72, p < .001, adjusted R-squared = .470).'
    )

    # TABLE 5: Regression Output
    add_table_with_data(doc,
        ['Model', 'R', 'R-squared', 'Adj R-squared', 'SE', 'F Change', 'Sig F Change'],
        [
            ['1 (Demographics)', '.287', '.082', '.080', '0.891', '36.94', '<.001'],
            ['2 (+Remote Work)', '.412', '.170', '.166', '0.848', '131.25', '<.001'],
            ['3 (+Work Environment)', '.583', '.340', '.336', '0.757', '160.83', '<.001'],
            ['4 (+Psychosocial)', '.688', '.473', '.470', '0.676', '158.72', '<.001'],
        ]
    )
    add_caption_paragraph(doc, 5, 'Regression Output')
    doc.add_paragraph('')

    doc.add_page_break()

    # ==================== Discussion ====================
    doc.add_heading('4. Discussion', level=1)
    doc.add_paragraph(
        'The findings from this comprehensive study paint a nuanced picture of remote work\'s impact on '
        'employee outcomes. The overall productivity gains observed (13.2% improvement) align with and '
        'extend previous research findings from pre-pandemic studies (Bloom et al., 2015; Choudhury et '
        'al., 2021), while our analysis of moderating factors provides new insights into the conditions '
        'under which remote work is most beneficial.'
    )

    # TABLE 6: Summary Statistics
    add_table_with_data(doc,
        ['Metric', 'Overall', 'Tech Sector', 'Finance', 'Healthcare', 'Education'],
        [
            ['Productivity Gain (%)', '13.2', '18.7', '14.1', '8.3', '9.6'],
            ['Satisfaction Increase', '0.24', '0.31', '0.28', '0.15', '0.19'],
            ['Well-Being Change', '+3.5', '+5.2', '+4.1', '+1.8', '+2.3'],
            ['Isolation Score', '3.56', '3.21', '3.45', '3.89', '3.78'],
            ['Burnout Decrease', '0.25', '0.34', '0.29', '0.12', '0.18'],
            ['Retention Intent (%)', '82.4', '88.1', '84.3', '76.5', '79.2'],
            ['Hybrid Preference (%)', '71.3', '78.4', '73.2', '65.8', '67.1'],
        ]
    )
    add_caption_paragraph(doc, 6, 'Summary Statistics')
    doc.add_paragraph('')

    doc.add_paragraph(
        'The significant role of managerial support as a moderating factor deserves particular attention. '
        'Organizations that invested in training managers for remote team leadership saw substantially '
        'better outcomes across all measured dimensions. This finding has direct practical implications '
        'for organizations designing hybrid work policies.'
    )

    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        'This study provides robust evidence that remote work, when properly supported, can enhance '
        'both productivity and employee well-being. However, the benefits are neither automatic nor '
        'universal. Organizational investment in managerial training, home office support, and social '
        'connection initiatives are critical moderators of success. Future research should examine '
        'longitudinal patterns as remote work becomes a permanent feature of the organizational landscape.'
    )

    doc.add_heading('References', level=1)
    doc.add_paragraph(
        'Bailey, D. E., & Kurland, N. B. (2002). A review of telework research: Findings, new '
        'directions, and lessons for the study of modern work. Journal of Organizational Behavior, '
        '23(4), 383-400.'
    )
    doc.add_paragraph(
        'Bloom, N., Liang, J., Roberts, J., & Ying, Z. J. (2015). Does working from home work? '
        'Evidence from a Chinese experiment. Quarterly Journal of Economics, 130(1), 165-218.'
    )
    doc.add_paragraph(
        'Choudhury, P., Foroughi, C., & Larson, B. (2021). Work-from-anywhere: The productivity '
        'effects of geographic flexibility. Strategic Management Journal, 42(4), 655-683.'
    )
    doc.add_paragraph(
        'Gajendran, R. S., & Harrison, D. A. (2007). The good, the bad, and the unknown about '
        'telecommuting: Meta-analysis of psychological mediators and individual consequences. '
        'Journal of Applied Psychology, 92(6), 1524-1541.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
