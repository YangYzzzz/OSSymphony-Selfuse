"""
Reward Script: Insert a Table of Tables that catalogs all numbered tables
Task ID: writer_mt_057
Domain: libreoffice_writer
Scoring:
  Component 1 (0.30) - "Table of Tables" heading exists (Heading 1 style)
  Component 2 (0.35) - TOC field code for Table category present
  Component 3 (0.35) - All 6 table captions listed as entries
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_mt_057'

# Expected table captions in the Table of Tables
EXPECTED_CAPTIONS = [
    "Table 1: Demographics",
    "Table 2: Survey Results",
    "Table 3: Statistical Analysis",
    "Table 4: Correlation Matrix",
    "Table 5: Regression Output",
    "Table 6: Summary Statistics",
]


def persist_app_state(domain):
    """Save any unsaved GUI edits before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: "Table of Tables" heading exists with Heading 1 style (0.30 points)
    # This heading is NOT present in the initial document (initial para[6] is 'Abstract').
    try:
        tot_heading_found = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            style_name = para.style.name if para.style else ""
            # Check for a heading-style paragraph containing "table of tables" or similar
            if "table" in text_lower and ("table of" in text_lower or "list of table" in text_lower):
                if "heading" in style_name.lower():
                    tot_heading_found = True
                    print(f"PASS: Component 1 - Found Table of Tables heading: '{para.text}' (style={style_name}) (0.30 pts)")
                    total_score += 0.30
                    break
        if not tot_heading_found:
            print("FAIL: Component 1 - No 'Table of Tables' heading found with Heading style")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: TOC field code for Table category present (0.35 points)
    # The golden document has instrText ' TOC \c "Table" ' which is the field code
    # for a Table of Figures index filtered to category "Table".
    # The initial document does NOT have this field code.
    try:
        body = doc.element.body
        toc_field_found = False
        for instr in body.findall('.//w:instrText', ns):
            instr_text = instr.text or ""
            # Look for TOC field code with Table category
            if "TOC" in instr_text and "Table" in instr_text:
                toc_field_found = True
                print(f"PASS: Component 2 - TOC field code found: {instr_text!r} (0.35 pts)")
                total_score += 0.35
                break
        if not toc_field_found:
            print("FAIL: Component 2 - No TOC field code for Table category found")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: All 6 table captions listed in the Table of Tables entries (0.35 points)
    # In the golden document, after the "Table of Tables" heading, there are 6 paragraphs
    # containing the table caption entries with page numbers. These paragraphs exist
    # only in the golden document, not in the initial.
    try:
        # Collect all paragraph texts to search for entries
        all_texts = [p.text.strip() for p in doc.paragraphs]

        # Find index of the Table of Tables heading
        tot_idx = -1
        for i, p in enumerate(doc.paragraphs):
            text_lower = p.text.strip().lower()
            style_name = p.style.name if p.style else ""
            if "table" in text_lower and ("table of" in text_lower or "list of table" in text_lower):
                if "heading" in style_name.lower():
                    tot_idx = i
                    break

        if tot_idx < 0:
            print("FAIL: Component 3 - Cannot find Table of Tables heading to locate entries")
        else:
            # Search paragraphs after the heading for table caption entries
            # Each entry should contain the caption text (e.g., "Table 1: Demographics")
            found_captions = 0
            # Search within a reasonable range after the heading
            search_range = doc.paragraphs[tot_idx + 1: tot_idx + 20]
            for caption in EXPECTED_CAPTIONS:
                caption_found = False
                for p in search_range:
                    p_text = p.text.strip()
                    # The entry text may have page number appended (e.g., "Table 1: Demographics\t5")
                    if caption in p_text:
                        caption_found = True
                        break
                if caption_found:
                    found_captions += 1

            # Award proportional points based on how many captions are found
            caption_ratio = found_captions / len(EXPECTED_CAPTIONS)
            points = round(0.35 * caption_ratio, 4)
            if found_captions == len(EXPECTED_CAPTIONS):
                print(f"PASS: Component 3 - All {found_captions}/{len(EXPECTED_CAPTIONS)} table captions found in entries ({points} pts)")
            else:
                print(f"PARTIAL: Component 3 - {found_captions}/{len(EXPECTED_CAPTIONS)} table captions found ({points} pts)")
            total_score += points
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    final_score = round(min(total_score, 1.0), 2)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
