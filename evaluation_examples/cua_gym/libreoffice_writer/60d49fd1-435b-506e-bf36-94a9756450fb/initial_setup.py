"""
Initial Setup: Insert page break before Appendix with page number restart
Task ID: wrpara_020
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'wrpara_020'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_page_number_footer(section):
    """Add a PAGE field code to the footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_prefix = fp.add_run("Page ")
    run_prefix.font.size = Pt(10)

    # PAGE field: begin
    r1 = fp.add_run()
    r1.font.size = Pt(10)
    fld_begin = r1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    r1._element.append(fld_begin)

    # PAGE field: instrText
    r2 = fp.add_run()
    r2.font.size = Pt(10)
    instr = r2._element.makeelement(qn('w:instrText'), {})
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)

    # PAGE field: end
    r3 = fp.add_run()
    r3.font.size = Pt(10)
    fld_end = r3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    r3._element.append(fld_end)


def create_initial():
    doc = Document()

    # Set default page layout
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Add page number footer to the single section
    add_page_number_footer(section)

    # ============================================================
    # SECTION 1: Introduction (~2 pages of content)
    # ============================================================
    h1 = doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'This technical report presents a comprehensive analysis of the automated '
        'quality assurance framework developed for Meridian Technologies Inc. during '
        'the fiscal year 2025. The framework was designed to address persistent challenges '
        'in software deployment pipelines, reduce defect rates in production releases, '
        'and establish standardized testing protocols across all development teams.'
    )

    doc.add_paragraph(
        'The initiative was launched in January 2025 following an internal audit that '
        'revealed a 23% increase in post-deployment defects compared to the previous year. '
        'The audit, conducted by the Quality Engineering division under the leadership of '
        'Dr. Rebecca Torres, identified three primary areas of concern: insufficient test '
        'coverage in microservices architectures, lack of automated regression testing for '
        'legacy components, and inconsistent code review practices across regional teams.'
    )

    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph(
        'Meridian Technologies operates a distributed software development infrastructure '
        'spanning four continents, with primary engineering hubs in Seattle, Dublin, '
        'Singapore, and Sao Paulo. The company maintains over 340 active software products '
        'serving approximately 12 million enterprise users. Each product undergoes an '
        'average of 47 deployments per quarter, resulting in over 16,000 annual deployment '
        'events across the organization.'
    )

    doc.add_paragraph(
        'Prior to this initiative, quality assurance processes were largely decentralized, '
        'with each regional team maintaining its own testing frameworks, deployment checklists, '
        'and defect tracking methodologies. While this approach allowed for regional flexibility, '
        'it created significant inconsistencies in quality outcomes. The Seattle hub reported '
        'an average defect density of 2.3 per KLOC, while the Singapore hub reported 4.1 per '
        'KLOC for comparable product lines.'
    )

    doc.add_heading('1.2 Objectives', level=2)
    doc.add_paragraph(
        'The primary objectives of the Automated Quality Assurance Framework project were '
        'established during the January 2025 planning summit attended by senior engineering '
        'leadership from all four regional hubs. The following goals were ratified:'
    )

    objectives = [
        'Reduce post-deployment defect rates by a minimum of 40% within 12 months',
        'Achieve 85% automated test coverage across all Tier-1 products by Q3 2025',
        'Implement continuous integration and continuous deployment (CI/CD) pipelines '
        'with integrated quality gates for all active product lines',
        'Establish a unified defect classification taxonomy and reporting framework',
        'Deploy automated performance regression testing for all customer-facing APIs',
        'Create a centralized dashboard for real-time quality metrics visibility',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_paragraph(
        'These objectives were aligned with Meridian Technologies\' broader strategic vision '
        'of achieving ISO 25010 compliance across its entire product portfolio by the end of '
        '2026. The framework was also expected to support the company\'s ongoing migration '
        'to a cloud-native architecture, which required fundamentally different testing '
        'approaches compared to the legacy monolithic systems.'
    )

    doc.add_paragraph(
        'The budget allocation for the initiative was $4.7 million, distributed across '
        'personnel costs ($2.8M), tooling and infrastructure ($1.2M), training programs '
        '($0.4M), and contingency reserves ($0.3M). A cross-functional steering committee '
        'was established with quarterly review checkpoints to monitor progress against '
        'key performance indicators.'
    )

    # ============================================================
    # SECTION 2: Methodology (~2 pages of content)
    # ============================================================
    doc.add_heading('Methodology', level=1)

    doc.add_paragraph(
        'The development of the Automated Quality Assurance Framework followed a phased '
        'approach, with each phase building upon the deliverables of the previous one. '
        'The methodology was grounded in industry best practices from the IEEE 829 standard '
        'for software testing documentation and the ISTQB Advanced Level Test Manager syllabus.'
    )

    doc.add_heading('2.1 Phase One: Assessment and Planning (Q1 2025)', level=2)
    doc.add_paragraph(
        'The first phase involved a thorough assessment of existing quality assurance '
        'processes across all regional hubs. A team of 14 quality engineers conducted '
        'on-site evaluations over a six-week period, documenting current workflows, '
        'tool inventories, and team capabilities. The assessment covered 87 development '
        'teams and 215 active projects.'
    )

    doc.add_paragraph(
        'Key findings from the assessment phase included: (a) only 34% of teams had '
        'automated unit test suites with coverage exceeding 60%, (b) integration testing '
        'was performed manually in 62% of projects, (c) performance testing was conducted '
        'only during major release cycles rather than continuously, and (d) 71% of teams '
        'lacked formal test environment management procedures.'
    )

    doc.add_heading('2.2 Phase Two: Framework Design (Q2 2025)', level=2)
    doc.add_paragraph(
        'Based on the assessment findings, the framework design team developed a '
        'three-tier architecture comprising: a foundation layer for test infrastructure '
        'management, a middleware layer for test orchestration and scheduling, and a '
        'presentation layer for dashboards, reporting, and alerting. The design was '
        'reviewed and approved by the Architecture Review Board in April 2025.'
    )

    doc.add_paragraph(
        'The foundation layer was built on Kubernetes-based test clusters deployed '
        'in each regional data center. Each cluster was provisioned with dedicated '
        'compute resources: 128 vCPUs, 512 GB RAM, and 10 TB NVMe storage for test '
        'artifact management. The clusters utilized a custom container orchestration '
        'framework called TestRunner Pro, developed in-house by the Platform Engineering '
        'team under the direction of Principal Engineer Akira Tanaka.'
    )

    doc.add_heading('2.3 Phase Three: Implementation (Q3 2025)', level=2)
    doc.add_paragraph(
        'The implementation phase began with pilot deployments in the Seattle and Dublin '
        'hubs, targeting 12 high-priority Tier-1 products. Each product team was assigned '
        'a dedicated quality coach who facilitated the transition from manual to automated '
        'testing workflows. The rollout followed a structured onboarding protocol:'
    )

    steps = [
        'Week 1-2: Team assessment and customized training plan development',
        'Week 3-4: Tool installation, configuration, and initial test suite migration',
        'Week 5-6: Parallel operation with existing processes and validation',
        'Week 7-8: Full cutover to automated pipeline with monitoring',
        'Week 9-12: Optimization and advanced feature adoption',
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_paragraph(
        'During the implementation phase, the team encountered several unexpected '
        'challenges. The most significant was the incompatibility between TestRunner Pro '
        'and the legacy Oracle-based test data management system used by the Sao Paulo hub. '
        'This required a three-week remediation effort to develop a compatibility adapter, '
        'which delayed the Phase Three timeline by approximately 15 business days.'
    )

    doc.add_heading('2.4 Phase Four: Validation and Rollout (Q4 2025)', level=2)
    doc.add_paragraph(
        'The validation phase employed a rigorous statistical methodology to assess '
        'framework effectiveness. Defect data from the 12 pilot products was compared '
        'against a control group of 12 comparable products that had not yet adopted the '
        'framework. The analysis used a paired t-test with a significance level of 0.05 '
        'and achieved a statistical power of 0.92.'
    )

    doc.add_paragraph(
        'Results from the validation phase demonstrated a 47% reduction in post-deployment '
        'defects for pilot products compared to the control group (p < 0.001). Automated '
        'test coverage increased from an average of 34% to 88% across pilot teams. Mean '
        'time to detect defects decreased from 4.3 days to 0.7 days, and mean time to '
        'resolve decreased from 6.2 days to 2.1 days.'
    )

    doc.add_paragraph(
        'Based on these results, the steering committee approved full organizational '
        'rollout beginning in January 2026. The rollout plan targets complete adoption '
        'across all 87 development teams by the end of Q2 2026, with provisional support '
        'structures maintained through Q4 2026 to ensure sustained adoption and compliance.'
    )

    # ============================================================
    # SECTION 3: Appendix (~1 page of content)
    # ============================================================
    doc.add_heading('Appendix', level=1)

    doc.add_heading('A. Tool Inventory', level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['Tool Name', 'Category', 'License Type', 'Annual Cost']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    tools_data = [
        ['TestRunner Pro', 'Test Orchestration', 'Internal', '$0'],
        ['Selenium Grid', 'UI Testing', 'Open Source', '$0'],
        ['JMeter Enterprise', 'Performance Testing', 'Commercial', '$45,000'],
        ['SonarQube Enterprise', 'Code Quality', 'Commercial', '$32,000'],
        ['Grafana Cloud', 'Monitoring', 'SaaS', '$18,500'],
        ['Artifactory Pro', 'Artifact Management', 'Commercial', '$27,000'],
        ['PagerDuty', 'Incident Management', 'SaaS', '$15,200'],
        ['Confluence', 'Documentation', 'SaaS', '$8,400'],
    ]
    for row_data in tools_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph('')  # spacer

    doc.add_heading('B. Key Personnel', level=2)
    doc.add_paragraph(
        'Dr. Rebecca Torres - VP of Quality Engineering (Project Sponsor)\n'
        'Akira Tanaka - Principal Engineer, Platform Engineering (Technical Lead)\n'
        'Sofia Martinez - Senior QA Manager, Seattle Hub (Implementation Lead)\n'
        'James O\'Brien - QA Director, Dublin Hub (Regional Coordinator)\n'
        'Priya Sharma - Test Architect, Singapore Hub (Regional Coordinator)\n'
        'Carlos Mendez - QA Manager, Sao Paulo Hub (Regional Coordinator)'
    )

    doc.add_heading('C. Glossary', level=2)
    glossary_items = [
        ('CI/CD', 'Continuous Integration / Continuous Deployment'),
        ('KLOC', 'Thousands of Lines of Code'),
        ('ISTQB', 'International Software Testing Qualifications Board'),
        ('NVMe', 'Non-Volatile Memory Express'),
        ('SaaS', 'Software as a Service'),
        ('vCPU', 'Virtual Central Processing Unit'),
    ]
    for term, definition in glossary_items:
        p = doc.add_paragraph()
        run_term = p.add_run(f'{term}: ')
        run_term.bold = True
        p.add_run(definition)

    # Save
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
