"""
Reward Script: Mail merge field insertion in Writer document
Task ID: writer_rd_049
Domain: libreoffice_writer
Scoring:
  Component 1 (0.25): MERGEFIELD FirstName exists in document
  Component 2 (0.25): MERGEFIELD LastName exists in document
  Component 3 (0.25): MERGEFIELD Company exists in document
  Component 4 (0.25): MERGEFIELD Address exists in document
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_049'


def persist_app_state(domain: str):
    """Save any unsaved changes in LibreOffice before verification."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def extract_merge_fields(doc):
    """Extract all MERGEFIELD names from the document XML."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    fields = []
    for para in doc.paragraphs:
        instr_elements = para._element.findall('.//w:instrText', ns)
        for instr in instr_elements:
            text = (instr.text or '').strip()
            if text.startswith('MERGEFIELD'):
                # Extract field name: "MERGEFIELD FirstName" -> "FirstName"
                parts = text.split()
                if len(parts) >= 2:
                    fields.append(parts[1])
    return fields


def check_plain_placeholders(doc):
    """Check if old plain-text placeholders still exist."""
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    placeholders = {
        'FirstName': '[First Name]',
        'LastName': '[Last Name]',
        'Company': '[Company]',
        'Address': '[Address]',
    }
    remaining = {}
    for field, placeholder in placeholders.items():
        if placeholder in full_text:
            remaining[field] = placeholder
    return remaining


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0

    The task requires replacing plain text placeholders with MERGEFIELD codes:
      [First Name] -> MERGEFIELD FirstName
      [Last Name]  -> MERGEFIELD LastName
      [Company]    -> MERGEFIELD Company
      [Address]    -> MERGEFIELD Address
    """
    from docx import Document
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract all merge fields from document
    merge_fields = extract_merge_fields(doc)
    print(f"INFO: Found merge fields: {merge_fields}")

    # Check which plain placeholders remain
    remaining_placeholders = check_plain_placeholders(doc)
    if remaining_placeholders:
        print(f"INFO: Plain text placeholders still present: {remaining_placeholders}")

    # Required merge fields and their expected minimum occurrences
    required_fields = {
        'FirstName': 1,   # appears in address block and salutation
        'LastName': 1,    # appears in address block and salutation
        'Company': 1,     # appears in address block and body
        'Address': 1,     # appears in address block
    }

    # Component 1: MERGEFIELD FirstName (0.25 points)
    try:
        fn_count = merge_fields.count('FirstName')
        if fn_count >= required_fields['FirstName'] and 'FirstName' not in remaining_placeholders:
            print(f"PASS: Component 1 — MERGEFIELD FirstName found {fn_count} time(s) (0.25 pts)")
            total_score += 0.25
        else:
            if fn_count == 0:
                print(f"FAIL: Component 1 — No MERGEFIELD FirstName found")
            else:
                print(f"FAIL: Component 1 — MERGEFIELD FirstName found but plain placeholder still exists")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: MERGEFIELD LastName (0.25 points)
    try:
        ln_count = merge_fields.count('LastName')
        if ln_count >= required_fields['LastName'] and 'LastName' not in remaining_placeholders:
            print(f"PASS: Component 2 — MERGEFIELD LastName found {ln_count} time(s) (0.25 pts)")
            total_score += 0.25
        else:
            if ln_count == 0:
                print(f"FAIL: Component 2 — No MERGEFIELD LastName found")
            else:
                print(f"FAIL: Component 2 — MERGEFIELD LastName found but plain placeholder still exists")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: MERGEFIELD Company (0.25 points)
    try:
        co_count = merge_fields.count('Company')
        if co_count >= required_fields['Company'] and 'Company' not in remaining_placeholders:
            print(f"PASS: Component 3 — MERGEFIELD Company found {co_count} time(s) (0.25 pts)")
            total_score += 0.25
        else:
            if co_count == 0:
                print(f"FAIL: Component 3 — No MERGEFIELD Company found")
            else:
                print(f"FAIL: Component 3 — MERGEFIELD Company found but plain placeholder still exists")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: MERGEFIELD Address (0.25 points)
    try:
        addr_count = merge_fields.count('Address')
        if addr_count >= required_fields['Address'] and 'Address' not in remaining_placeholders:
            print(f"PASS: Component 4 — MERGEFIELD Address found {addr_count} time(s) (0.25 pts)")
            total_score += 0.25
        else:
            if addr_count == 0:
                print(f"FAIL: Component 4 — No MERGEFIELD Address found")
            else:
                print(f"FAIL: Component 4 — MERGEFIELD Address found but plain placeholder still exists")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persist any unsaved edits before scoring
persist_app_state("libreoffice_writer")

# Test against canonical artifact path
file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
