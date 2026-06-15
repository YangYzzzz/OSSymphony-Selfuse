"""
Initial Setup: Mail merge letter template with plain text placeholders
Task ID: writer_rd_049
Domain: libreoffice_writer
"""

import csv
import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rd_049'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'
CSV_PATH = f'{WORKDIR}/clients.csv'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_csv():
    """Create the clients.csv data source with 25 records."""
    clients = [
        ("Sarah", "Chen", "Nexus Innovations", "1240 Market Street", "San Francisco", "CA", "94102"),
        ("Marcus", "Johnson", "Apex Dynamics", "875 N Michigan Ave", "Chicago", "IL", "60611"),
        ("Elena", "Rodriguez", "Pinnacle Systems", "200 Congress Ave", "Austin", "TX", "78701"),
        ("David", "Kim", "Summit Analytics", "1500 Broadway", "New York", "NY", "10036"),
        ("Priya", "Patel", "Horizon Technologies", "3400 Hillview Ave", "Palo Alto", "CA", "94304"),
        ("James", "O'Brien", "Vanguard Solutions", "100 Federal Street", "Boston", "MA", "02110"),
        ("Mei", "Zhang", "Catalyst Group", "600 Anton Blvd", "Costa Mesa", "CA", "92626"),
        ("Robert", "Williams", "Sterling Partners", "2100 Ross Ave", "Dallas", "TX", "75201"),
        ("Aisha", "Hassan", "Quantum Ventures", "1201 Third Ave", "Seattle", "WA", "98101"),
        ("Thomas", "Mueller", "Alpine Consulting", "333 S Wabash Ave", "Chicago", "IL", "60604"),
        ("Linda", "Nakamura", "Pacific Rim Advisors", "700 Bishop Street", "Honolulu", "HI", "96813"),
        ("Carlos", "Mendez", "Global Edge Corp", "1401 Brickell Ave", "Miami", "FL", "33131"),
        ("Jennifer", "Walsh", "Beacon Strategies", "250 E Ponce de Leon Ave", "Decatur", "GA", "30030"),
        ("Rajesh", "Gupta", "Vertex Dynamics", "1001 Woodward Ave", "Detroit", "MI", "48226"),
        ("Sofia", "Petrov", "Emerald Analytics", "1420 Fifth Ave", "Seattle", "WA", "98101"),
        ("Michael", "Thompson", "Redwood Enterprises", "1 Embarcadero Center", "San Francisco", "CA", "94111"),
        ("Fatima", "Al-Rashid", "Oasis International", "301 Commerce Street", "Nashville", "TN", "37201"),
        ("Patrick", "O'Connor", "Celtic Innovations", "150 S Independence Mall", "Philadelphia", "PA", "19106"),
        ("Yuki", "Tanaka", "Sakura Dynamics", "1 World Trade Center", "New York", "NY", "10007"),
        ("Amanda", "Foster", "Brightpath Solutions", "200 E Colfax Ave", "Denver", "CO", "80203"),
        ("Victor", "Ivanov", "Northern Lights Tech", "1100 S Washington St", "Minneapolis", "MN", "55415"),
        ("Grace", "Lee", "Harmony Consulting", "100 N Tryon Street", "Charlotte", "NC", "28202"),
        ("Omar", "Sayed", "Crescent Holdings", "600 Travis Street", "Houston", "TX", "77002"),
        ("Catherine", "Dubois", "Fleur Analytics", "1 Canal Street", "New Orleans", "LA", "70130"),
        ("Daniel", "Park", "Aspen Group", "1625 Broadway", "Denver", "CO", "80202"),
    ]

    with open(CSV_PATH, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(["FirstName", "LastName", "Company", "Address", "City", "State", "Zip"])
        for c in clients:
            writer.writerow(c)

    print(f'CSV data source created: {CSV_PATH} ({len(clients)} records)')


def create_initial_document():
    """Create business letter template with plain text placeholders."""
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # --- Company letterhead ---
    letterhead = doc.add_paragraph()
    letterhead.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    letterhead.paragraph_format.space_after = Pt(0)
    run = letterhead.add_run("Meridian Consulting Group")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    addr_line = doc.add_paragraph()
    addr_line.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr_line.paragraph_format.space_before = Pt(2)
    addr_line.paragraph_format.space_after = Pt(0)
    run = addr_line.add_run("4500 Executive Parkway, Suite 300  |  Portland, OR 97201")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    contact_line = doc.add_paragraph()
    contact_line.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_line.paragraph_format.space_before = Pt(0)
    contact_line.paragraph_format.space_after = Pt(6)
    run = contact_line.add_run("Tel: (503) 555-0184  |  info@meridianconsulting.com  |  www.meridianconsulting.com")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Horizontal rule (thin line) ---
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(0)
    hr_para.paragraph_format.space_after = Pt(12)
    pBdr = hr_para._element.get_or_add_pPr()
    from docx.oxml.ns import qn
    bdr = pBdr.makeelement(qn('w:pBdr'), {})
    bottom = bdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '1F497D',
    })
    bdr.append(bottom)
    pBdr.append(bdr)

    # --- Date ---
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(18)
    run = date_para.add_run("March 28, 2026")
    run.font.size = Pt(11)

    # --- Recipient address block with placeholders ---
    recipient_lines = [
        "[First Name] [Last Name]",
        "[Company]",
        "[Address]",
    ]
    for line in recipient_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)

    # Spacing after address block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(12)
    spacer.paragraph_format.space_after = Pt(0)

    # --- Salutation ---
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(6)
    run = salutation.add_run("Dear [First Name] [Last Name],")
    run.font.size = Pt(11)

    # --- Body paragraphs ---
    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after = Pt(6)
    run = body1.add_run(
        "Thank you for your continued partnership with Meridian Consulting Group. "
        "We are writing to inform you about our upcoming Strategic Planning Workshop "
        "scheduled for May 15-17, 2026, at the Portland Convention Center."
    )
    run.font.size = Pt(11)

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after = Pt(6)
    run = body2.add_run(
        "As a valued client of [Company], we believe this workshop will provide "
        "exceptional value for your team. The three-day intensive program covers "
        "market analysis methodologies, competitive positioning strategies, and "
        "organizational transformation frameworks that have driven measurable results "
        "for over 200 organizations nationwide."
    )
    run.font.size = Pt(11)

    body3 = doc.add_paragraph()
    body3.paragraph_format.space_after = Pt(6)
    run = body3.add_run(
        "Early registration is now open, and we are pleased to offer our existing "
        "clients a 20% discount on the standard registration fee. Seating is limited "
        "to 50 participants to ensure an interactive and personalized experience."
    )
    run.font.size = Pt(11)

    body4 = doc.add_paragraph()
    body4.paragraph_format.space_after = Pt(6)
    run = body4.add_run(
        "Please do not hesitate to reach out if you have any questions or would like "
        "to reserve your spot. We look forward to welcoming you at the workshop."
    )
    run.font.size = Pt(11)

    # --- Closing ---
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(0)
    run = closing.add_run("Warm regards,")
    run.font.size = Pt(11)

    # Signature block
    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_before = Pt(24)
    sig_name.paragraph_format.space_after = Pt(0)
    run = sig_name.add_run("Victoria Hargrove")
    run.bold = True
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_before = Pt(0)
    sig_title.paragraph_format.space_after = Pt(0)
    run = sig_title.add_run("Senior Partner & Workshop Director")
    run.font.size = Pt(11)

    sig_co = doc.add_paragraph()
    sig_co.paragraph_format.space_before = Pt(0)
    sig_co.paragraph_format.space_after = Pt(0)
    run = sig_co.add_run("Meridian Consulting Group")
    run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f'Initial document created: {OUTPUT}')


def main():
    create_csv()
    create_initial_document()

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


main()
