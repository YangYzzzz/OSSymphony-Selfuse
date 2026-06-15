"""
Reward Script: Configure different left/right page headers in thesis
Task ID: writer_acad_052
Domain: libreoffice_writer
Scoring:
  Component 1: evenAndOddHeaders enabled (0.25)
  Component 2: Even (left) page header contains thesis title (0.25)
  Component 3: Odd (right) page header contains chapter title (0.25)
  Component 4: Header alignments correct (left for even, right for odd) (0.25)
"""

import os

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_052'

THESIS_TITLE = 'Machine Learning for Climate Science'


def persist_app_state(domain):
    """Best-effort save of any open LibreOffice document."""
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print(f"PERSIST: ctrl+s sent for {domain}")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        print(f"CRITICAL: Cannot import docx: {e}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: evenAndOddHeaders is enabled in document settings (0.25 points)
    # This is the core setting that enables different left/right page headers.
    # In the initial file this element is absent; in the golden file it is present.
    try:
        settings_elem = doc.settings._element
        eaoh = settings_elem.find(qn('w:evenAndOddHeaders'))
        if eaoh is not None:
            # Check that it is not explicitly set to false
            val = eaoh.get(qn('w:val'))
            # val=None or val="1" or val="true" all mean enabled
            if val is None or val in ('1', 'true'):
                print("PASS: Component 1 -- evenAndOddHeaders is enabled (0.25 pts)")
                total_score += 0.25
            else:
                print(f"FAIL: Component 1 -- evenAndOddHeaders val={val} (expected enabled)")
        else:
            print("FAIL: Component 1 -- evenAndOddHeaders element not found in settings")
    except Exception as e:
        print(f"ERROR: Component 1 -- {e}")

    # Component 2: Even (left) page header contains thesis title (0.25 points)
    # The even page header should contain 'Machine Learning for Climate Science'.
    # In the initial file, the even header is linked_to_previous and empty.
    try:
        section = doc.sections[0]
        even_header = section.even_page_header
        even_linked = even_header.is_linked_to_previous
        even_text = ''
        for p in even_header.paragraphs:
            if p.text.strip():
                even_text = p.text.strip()
                break

        if even_linked:
            print("FAIL: Component 2 -- Even header is linked to previous (no separate even header)")
        elif THESIS_TITLE.lower() in even_text.lower():
            print(f"PASS: Component 2 -- Even header contains thesis title: '{even_text}' (0.25 pts)")
            total_score += 0.25
        else:
            print(f"FAIL: Component 2 -- Even header text is '{even_text}', expected '{THESIS_TITLE}'")
    except Exception as e:
        print(f"ERROR: Component 2 -- {e}")

    # Component 3: Odd (right/default) page header contains a chapter title, not the thesis title (0.25 points)
    # The default/odd header should contain a chapter title text (not the thesis title).
    # In the initial file, the odd header contains the thesis title (same on all pages).
    try:
        section = doc.sections[0]
        header = section.header
        odd_text = ''
        for p in header.paragraphs:
            if p.text.strip():
                odd_text = p.text.strip()
                break

        # Odd header must have non-empty text that differs from the thesis title
        has_chapter_text = (odd_text and THESIS_TITLE.lower() != odd_text.lower())
        if has_chapter_text:
            print(f"PASS: Component 3 -- Odd header contains chapter title: '{odd_text}' (0.25 pts)")
            total_score += 0.25
        elif not odd_text:
            print("FAIL: Component 3 -- Odd header is empty")
        else:
            print(f"FAIL: Component 3 -- Odd header still has thesis title '{odd_text}', expected chapter title")
    except Exception as e:
        print(f"ERROR: Component 3 -- {e}")

    # Component 4: Header alignments (0.25 points)
    # Even (left page) header should be LEFT aligned, odd (right page) header should be RIGHT aligned.
    # In the initial file, the single header is CENTER aligned.
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        section = doc.sections[0]

        # Check even header alignment
        even_header = section.even_page_header
        even_align = None
        for p in even_header.paragraphs:
            if p.text.strip():
                even_align = p.paragraph_format.alignment
                break

        # Check odd/default header alignment
        odd_header = section.header
        odd_align = None
        for p in odd_header.paragraphs:
            if p.text.strip():
                odd_align = p.paragraph_format.alignment
                break

        even_ok = (even_align is not None and
                   even_align == WD_PARAGRAPH_ALIGNMENT.LEFT)
        odd_ok = (odd_align is not None and
                  odd_align == WD_PARAGRAPH_ALIGNMENT.RIGHT)

        if even_ok and odd_ok:
            print(f"PASS: Component 4 -- Even header LEFT aligned, odd header RIGHT aligned (0.25 pts)")
            total_score += 0.25
        else:
            reasons = []
            if not even_ok:
                reasons.append(f"even alignment={even_align} (expected LEFT)")
            if not odd_ok:
                reasons.append(f"odd alignment={odd_align} (expected RIGHT)")
            print(f"FAIL: Component 4 -- {'; '.join(reasons)}")
    except Exception as e:
        print(f"ERROR: Component 4 -- {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state('libreoffice_writer')

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
