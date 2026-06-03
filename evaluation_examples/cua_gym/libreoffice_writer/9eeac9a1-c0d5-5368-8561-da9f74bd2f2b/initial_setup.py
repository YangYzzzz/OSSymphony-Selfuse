"""
Initial Setup: Thesis document with same header on all pages
Task ID: writer_acad_052
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_052'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_heading_styled(doc, text, level=1):
    """Add a heading with explicit formatting."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_body_paragraph(doc, text, space_after=Pt(6)):
    """Add a body paragraph with standard thesis formatting."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = space_after
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    return para


def create_initial():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # -- Set a UNIFORM header on all pages (same on left and right) --
    # This is the default behavior: same header on all pages
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Machine Learning for Climate Science"
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in hp.runs:
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.italic = True

    # -- Title Page --
    title_para = doc.add_paragraph()
    title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    run = title_para.add_run("Machine Learning for Climate Science")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    run = subtitle.add_run("A Thesis Submitted in Partial Fulfillment\nof the Requirements for the Degree of\nDoctor of Philosophy")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    author = doc.add_paragraph()
    author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_before = Pt(48)
    run = author.add_run("Elena Rodriguez-Vasquez")
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    dept = doc.add_paragraph()
    dept.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    dept.paragraph_format.space_before = Pt(12)
    run = dept.add_run("Department of Earth and Atmospheric Sciences\nStanford University\nMarch 2025")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Page break after title page
    doc.add_page_break()

    # -- Chapter 1: Introduction --
    add_heading_styled(doc, "Chapter 1: Introduction", level=1)

    add_body_paragraph(doc,
        "Climate science has entered a transformative era driven by the exponential growth "
        "of observational data and advances in computational methods. The ability to model "
        "and predict climate phenomena with increasing accuracy has profound implications "
        "for policy-making, disaster preparedness, and sustainable development.")

    add_body_paragraph(doc,
        "Traditional climate models rely on numerical solutions to coupled partial "
        "differential equations governing atmospheric and oceanic circulation. While these "
        "physics-based models have achieved remarkable success, they face fundamental "
        "limitations in resolution, parameterization of sub-grid processes, and "
        "computational cost. A single high-resolution climate simulation can require "
        "millions of CPU hours on supercomputing facilities.")

    add_body_paragraph(doc,
        "Machine learning offers a complementary approach that can address several of "
        "these limitations. Neural networks excel at learning complex nonlinear mappings "
        "from data, and recent architectures such as transformers and graph neural networks "
        "have demonstrated remarkable ability to capture spatiotemporal dependencies in "
        "geophysical data. This thesis investigates the application of modern machine "
        "learning techniques to three critical challenges in climate science.")

    add_body_paragraph(doc,
        "The first challenge concerns precipitation forecasting at mesoscale resolution. "
        "Precipitation remains one of the most difficult variables to predict accurately, "
        "owing to its intermittent nature and strong dependence on local topography and "
        "convective processes. We develop a hybrid approach combining convolutional neural "
        "networks with attention mechanisms to produce 6-hour precipitation forecasts at "
        "4 km resolution over the continental United States.")

    add_body_paragraph(doc,
        "The second challenge addresses the emulation of computationally expensive "
        "radiative transfer calculations. Radiation schemes account for approximately 30% "
        "of the total computational cost in global climate models. We train a deep neural "
        "network emulator that achieves radiative flux errors below 0.5 W/m\u00b2 while "
        "providing a 200x speedup over the reference scheme.")

    doc.add_page_break()

    # -- Chapter 2: Background and Related Work --
    add_heading_styled(doc, "Chapter 2: Background and Related Work", level=1)

    add_heading_styled(doc, "2.1 Climate Modeling Fundamentals", level=2)

    add_body_paragraph(doc,
        "The foundation of modern climate modeling rests on the primitive equations "
        "governing atmospheric motion: the Navier-Stokes equations under the hydrostatic "
        "approximation, the thermodynamic energy equation, the continuity equation for "
        "dry air, and conservation equations for water substance. These equations are "
        "discretized on computational grids and integrated forward in time using numerical "
        "methods.")

    add_body_paragraph(doc,
        "General Circulation Models (GCMs) solve these equations on a global grid with "
        "typical horizontal resolutions of 50\u2013200 km. At these resolutions, many "
        "important physical processes\u2014including cloud formation, convection, turbulent "
        "mixing, and precipitation\u2014occur at scales smaller than the grid spacing and "
        "must be represented through parameterization schemes. The choice and tuning of "
        "these parameterizations represents one of the largest sources of uncertainty in "
        "climate projections (Schneider et al., 2017).")

    add_heading_styled(doc, "2.2 Machine Learning in Earth Sciences", level=2)

    add_body_paragraph(doc,
        "The application of machine learning to earth sciences dates back to early work "
        "on neural network weather prediction in the 1990s (Hsieh and Tang, 1998). "
        "However, it was the deep learning revolution beginning around 2012 that catalyzed "
        "rapid growth in the field. Key developments include the use of convolutional "
        "neural networks for remote sensing image classification (Zhu et al., 2017), "
        "recurrent networks for time series forecasting (Shi et al., 2015), and generative "
        "models for statistical downscaling (Vandal et al., 2017).")

    add_body_paragraph(doc,
        "Recent years have seen the emergence of foundation models for weather prediction, "
        "including GraphCast (Lam et al., 2023), Pangu-Weather (Bi et al., 2023), and "
        "FourCastNet (Pathak et al., 2022). These models, trained on decades of reanalysis "
        "data, have demonstrated skill competitive with or exceeding operational numerical "
        "weather prediction systems for medium-range forecasts, while requiring orders of "
        "magnitude less computational time at inference.")

    add_heading_styled(doc, "2.3 Physics-Informed Neural Networks", level=2)

    add_body_paragraph(doc,
        "Physics-informed neural networks (PINNs) incorporate known physical laws directly "
        "into the loss function during training, constraining the learned representations "
        "to be physically consistent. Originally proposed by Raissi et al. (2019) for "
        "solving forward and inverse problems involving partial differential equations, "
        "PINNs have since been adapted for climate applications including ocean modeling "
        "(Mavi et al., 2023) and atmospheric chemistry (Kelp et al., 2022).")

    doc.add_page_break()

    # -- Chapter 3: Precipitation Forecasting --
    add_heading_styled(doc, "Chapter 3: Precipitation Forecasting with Hybrid Neural Networks", level=1)

    add_heading_styled(doc, "3.1 Data and Preprocessing", level=2)

    add_body_paragraph(doc,
        "We utilize the Multi-Radar Multi-Sensor (MRMS) precipitation dataset, which "
        "provides gauge-corrected radar-based precipitation estimates at 1 km resolution "
        "over the contiguous United States. Our study period spans January 2016 through "
        "December 2023, with data aggregated to 4 km resolution and 1-hour temporal "
        "granularity. Atmospheric state variables are obtained from the HRRR (High-Resolution "
        "Rapid Refresh) model analyses at matching resolution.")

    add_body_paragraph(doc,
        "Input features include 17 atmospheric variables at 5 pressure levels (surface, "
        "850, 700, 500, and 300 hPa): temperature, specific humidity, u- and v-wind "
        "components, geopotential height, vertical velocity, and relative humidity. "
        "Additionally, we include static features: elevation, land use classification, "
        "and latitude/longitude encodings. Total input dimensionality is 92 channels "
        "per grid point.")

    add_heading_styled(doc, "3.2 Model Architecture", level=2)

    add_body_paragraph(doc,
        "Our proposed PrecipNet architecture combines a U-Net backbone for multi-scale "
        "spatial feature extraction with a cross-attention module that captures long-range "
        "dependencies. The encoder pathway consists of four convolutional blocks, each "
        "with two 3x3 convolution layers followed by batch normalization, GELU activation, "
        "and 2x2 max pooling. The bottleneck layer incorporates a multi-head self-attention "
        "mechanism with 8 heads operating on flattened spatial features.")

    add_body_paragraph(doc,
        "The decoder pathway mirrors the encoder with transposed convolutions for "
        "upsampling and skip connections from corresponding encoder levels. A novel "
        "cross-attention block at each decoder level attends to features from the "
        "previous 6 time steps, enabling the model to track storm evolution and "
        "propagation. The final output layer produces probabilistic forecasts via a "
        "mixture density network with 3 Gaussian components, capturing the multimodal "
        "nature of precipitation distributions.")

    add_heading_styled(doc, "3.3 Results", level=2)

    add_body_paragraph(doc,
        "PrecipNet achieves a Critical Success Index (CSI) of 0.47 for heavy precipitation "
        "events (>10 mm/hr) at 6-hour lead time, compared to 0.38 for the HRRR model and "
        "0.41 for a baseline U-Net without attention. The improvement is most pronounced "
        "for organized convective systems, where the attention mechanism successfully "
        "captures storm propagation patterns. Continuous Ranked Probability Skill Score "
        "(CRPSS) averaged over all precipitation thresholds is 0.32, representing a 15% "
        "improvement over the deterministic HRRR baseline.")

    doc.add_page_break()

    # -- Chapter 4: Radiative Transfer Emulation --
    add_heading_styled(doc, "Chapter 4: Neural Network Emulation of Radiative Transfer", level=1)

    add_body_paragraph(doc,
        "Radiative transfer calculations in climate models compute the absorption, "
        "emission, and scattering of solar and terrestrial radiation as it passes through "
        "the atmosphere. The RRTMGP (Rapid Radiative Transfer Model for GCMs, Parallel) "
        "scheme used in many modern climate models performs these calculations across 224 "
        "spectral intervals for longwave radiation and 112 intervals for shortwave radiation. "
        "This process must be repeated for every grid column at every time step, making it "
        "one of the most computationally intensive components of climate simulations.")

    add_body_paragraph(doc,
        "We develop RadianceNet, a deep neural network emulator trained on input-output "
        "pairs generated by RRTMGP. The training dataset comprises 50 million atmospheric "
        "profiles sampled from 10 years of CESM2 climate simulations, spanning diverse "
        "conditions from polar to tropical latitudes, clear and cloudy skies, and day and "
        "night conditions. Each profile includes 60 vertical levels of temperature, "
        "pressure, humidity, ozone concentration, and cloud properties.")

    add_body_paragraph(doc,
        "RadianceNet employs a residual MLP architecture with 8 hidden layers of 512 "
        "units each, using skip connections every 2 layers. The model is trained to predict "
        "heating rates and radiative fluxes at each vertical level simultaneously. Online "
        "validation within the CESM2 framework shows global-mean top-of-atmosphere flux "
        "errors of 0.3 W/m\u00b2 for longwave and 0.4 W/m\u00b2 for shortwave, well within "
        "the intermodel spread of CMIP6 ensembles. The emulator achieves a 200x wall-clock "
        "speedup over RRTMGP while maintaining numerical stability in year-long coupled "
        "climate simulations.")

    # -- Save --
    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
