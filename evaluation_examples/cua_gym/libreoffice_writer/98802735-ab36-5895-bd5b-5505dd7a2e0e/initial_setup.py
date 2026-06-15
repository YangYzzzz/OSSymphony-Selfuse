"""
Initial Setup: Create a literature review document with author citations
Task ID: writer_acad_074
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_074'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # -- Title --
    title = doc.add_heading('Literature Review: Advances in Computational Neuroscience', level=1)

    # -- Introduction --
    doc.add_heading('1. Introduction', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'The field of computational neuroscience has undergone significant transformations '
        'over the past two decades. This literature review examines key contributions from '
        'leading researchers whose work has shaped our understanding of neural computation, '
        'synaptic plasticity, and brain-computer interfaces. The review is organized thematically, '
        'covering theoretical frameworks, experimental methodologies, and emerging applications.'
    )

    # -- Section 2: Theoretical Frameworks --
    doc.add_heading('2. Theoretical Frameworks', level=2)

    p = doc.add_paragraph()
    p.add_run(
        'Smith (2019) proposed a unified model of cortical microcircuit dynamics that integrates '
        'excitatory and inhibitory balance with attractor network theory. This framework, which '
        'built upon earlier work by Johnson (2017), provided a mathematical foundation for '
        'understanding how populations of neurons encode and process sensory information. '
        'Smith further demonstrated that the model could predict neural firing patterns in '
        'primary visual cortex with remarkable accuracy (Smith, 2021).'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Williams (2018) challenged several assumptions in the Smith framework by introducing '
        'a stochastic resonance component. Williams argued that noise in neural circuits is not '
        'merely a nuisance variable but plays a functional role in signal detection. This view '
        'was later supported by Brown (2020), who conducted extensive simulations showing that '
        'moderate noise levels enhanced information transmission in feed-forward networks.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Davis (2019) took a different approach by focusing on the temporal dynamics of '
        'neural coding. Rather than steady-state models, Davis emphasized transient responses '
        'and their role in rapid decision-making. The Davis temporal coding framework was '
        'subsequently validated through electrophysiological recordings in primates (Davis & '
        'Johnson, 2021).'
    )

    # -- Section 3: Experimental Methodologies --
    doc.add_heading('3. Experimental Methodologies', level=2)

    p = doc.add_paragraph()
    p.add_run(
        'Johnson (2018) developed a novel multi-electrode array technique that allowed '
        'simultaneous recording from over 500 neurons in freely moving animals. This '
        'methodological advance enabled researchers to study population dynamics at an '
        'unprecedented scale. Brown (2019) adapted the Johnson technique for use in '
        'cortical slice preparations, demonstrating that similar population-level patterns '
        'emerge in reduced preparations.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Williams and Davis (2020) combined optogenetic stimulation with calcium imaging '
        'to create a closed-loop system for studying causal relationships in neural circuits. '
        'Their approach allowed precise manipulation of specific cell types while monitoring '
        'network-wide responses. Smith (2020) later applied this methodology to investigate '
        'the role of parvalbumin-positive interneurons in gamma oscillations.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Brown (2021) introduced a machine learning pipeline for automated spike sorting '
        'that reduced manual curation time by 85%. The pipeline, validated against expert '
        'annotations from the Johnson laboratory, achieved classification accuracy exceeding '
        '97% across diverse recording conditions. Williams (2021) extended this approach to '
        'handle chronic recordings spanning several months.'
    )

    # -- Section 4: Applications and Translational Research --
    doc.add_heading('4. Applications and Translational Research', level=2)

    p = doc.add_paragraph()
    p.add_run(
        'The translational potential of computational neuroscience has been explored by '
        'several groups. Davis (2020) developed a brain-computer interface prototype that '
        'utilized the temporal coding framework for decoding motor intentions. Clinical '
        'trials conducted by Davis and Brown (2022) demonstrated that patients with spinal '
        'cord injuries could achieve cursor control within the first session of training.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Smith and Williams (2022) published a comprehensive review of neuroprosthetic '
        'applications, arguing that current devices underutilize the rich temporal structure '
        'of neural signals. Their analysis of 47 published studies revealed that incorporating '
        'temporal features improved decoding performance by an average of 23%. Johnson (2022) '
        'provided a complementary perspective, emphasizing the engineering challenges of '
        'real-time neural signal processing in implantable devices.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'Brown and Davis (2023) recently demonstrated a fully implantable neural recording '
        'system with on-chip spike sorting capabilities. This device, which incorporated the '
        'automated classification algorithms from Brown (2021), represented a significant '
        'step toward chronic, autonomous brain-computer interfaces. Williams (2023) noted that '
        'such advances bring closer the possibility of seamless neural prosthetics.'
    )

    # -- Section 5: Discussion --
    doc.add_heading('5. Discussion and Future Directions', level=2)

    p = doc.add_paragraph()
    p.add_run(
        'This review has highlighted the interconnected contributions of Smith, Johnson, '
        'Williams, Brown, and Davis to the field of computational neuroscience. Their work '
        'spans from theoretical models (Smith, 2019; Williams, 2018; Davis, 2019) through '
        'experimental innovations (Johnson, 2018; Brown, 2019) to translational applications '
        '(Davis, 2020; Brown & Davis, 2023). Future research should focus on integrating '
        'these diverse approaches into a cohesive framework for understanding brain function '
        'and developing next-generation neural technologies.'
    )

    # -- References --
    doc.add_heading('References', level=2)

    references = [
        'Brown, A. R. (2019). Population dynamics in cortical slice preparations using adapted multi-electrode arrays. Journal of Neuroscience Methods, 312, 45-58.',
        'Brown, A. R. (2020). Stochastic resonance enhances information transmission in feed-forward neural networks. Neural Computation, 32(8), 1567-1589.',
        'Brown, A. R. (2021). Automated spike sorting using deep learning: A validated pipeline for large-scale recordings. Nature Methods, 18(4), 412-420.',
        'Brown, A. R., & Davis, P. L. (2022). Brain-computer interface for cursor control in spinal cord injury patients. The Lancet Neurology, 21(3), 245-256.',
        'Brown, A. R., & Davis, P. L. (2023). Fully implantable neural recording system with on-chip classification. Science, 379(6631), 558-563.',
        'Davis, P. L. (2019). Temporal dynamics in neural coding: Beyond rate models. Trends in Neurosciences, 42(7), 501-513.',
        'Davis, P. L. (2020). A brain-computer interface based on temporal coding of motor intentions. IEEE Transactions on Biomedical Engineering, 67(9), 2534-2543.',
        'Davis, P. L., & Johnson, M. K. (2021). Validation of the temporal coding framework in primate electrophysiology. Proceedings of the National Academy of Sciences, 118(23), e2105678118.',
        'Johnson, M. K. (2017). Attractor network models of cortical computation. Annual Review of Neuroscience, 40, 289-312.',
        'Johnson, M. K. (2018). High-density multi-electrode recordings in freely moving animals. Nature Neuroscience, 21(6), 853-861.',
        'Johnson, M. K. (2022). Engineering challenges for real-time neural signal processing. Bioelectronic Medicine, 8(1), 15.',
        'Smith, R. J. (2019). A unified model of cortical microcircuit dynamics. Neuron, 103(4), 678-694.',
        'Smith, R. J. (2020). Parvalbumin interneurons and gamma oscillations: A closed-loop investigation. Cell Reports, 31(7), 107645.',
        'Smith, R. J. (2021). Predicting neural firing patterns in primary visual cortex. eLife, 10, e67890.',
        'Smith, R. J., & Williams, T. E. (2022). Neuroprosthetic applications: Underutilization of temporal neural features. Nature Reviews Neuroscience, 23(5), 301-316.',
        'Williams, T. E. (2018). Stochastic resonance in neural circuits: A functional perspective. PLoS Computational Biology, 14(6), e1006178.',
        'Williams, T. E. (2021). Chronic spike sorting across months-long recordings. Journal of Neural Engineering, 18(4), 046012.',
        'Williams, T. E. (2023). Toward seamless neural prosthetics: Current progress and remaining challenges. Nature Biomedical Engineering, 7(2), 145-158.',
        'Williams, T. E., & Davis, P. L. (2020). Closed-loop optogenetics with calcium imaging for causal circuit analysis. Nature Protocols, 15(9), 2891-2915.',
    ]

    for ref in references:
        p = doc.add_paragraph(ref, style='List Number')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
