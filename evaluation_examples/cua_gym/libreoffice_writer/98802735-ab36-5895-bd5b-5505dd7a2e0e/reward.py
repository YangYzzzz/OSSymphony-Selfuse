"""
Reward Script: Create an author index for a literature review chapter
Task ID: writer_acad_074
Domain: libreoffice_writer
Scoring:
  Component 1 (0.50): XE index entries exist for all 5 required authors
  Component 2 (0.20): INDEX field is present in the document
  Component 3 (0.15): Author Index heading/section exists at end of document
  Component 4 (0.15): Sufficient XE entry coverage (multiple occurrences marked)
"""

import os
import re
from collections import Counter

WORKDIR = '/home/user'
TASK_ID = 'writer_acad_074'
REQUIRED_AUTHORS = {'Smith', 'Johnson', 'Williams', 'Brown', 'Davis'}


def persist_app_state(domain: str):
    """Save any unsaved changes in LibreOffice Writer."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_writer", "libreoffice_calc", "libreoffice_impress"}:
        try:
            import pyautogui
            import time
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        from lxml import etree
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Get the full body XML for field searching
    try:
        body_xml = etree.tostring(doc.element.body, encoding='unicode')
    except Exception as e:
        print(f"CRITICAL: Cannot serialize XML: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Extract all XE entries from instrText elements
    xe_authors = []
    try:
        from docx.oxml.ns import qn
        for el in doc.element.body.iter(qn('w:instrText')):
            txt = el.text or ''
            # Match XE "AuthorName" pattern
            m = re.search(r'XE\s+.(\w+)', txt)
            if m:
                xe_authors.append(m.group(1))
    except Exception as e:
        print(f"WARN: XML iter failed, falling back to regex: {e}")
        # Fallback: regex on raw XML
        xe_authors_raw = re.findall(r'XE\s+["\u201c](\w+)', body_xml)
        xe_authors = xe_authors_raw

    xe_counter = Counter(xe_authors)
    found_authors = set(xe_counter.keys()) & REQUIRED_AUTHORS

    # Component 1: XE index entries exist for all 5 required authors (0.50 points)
    # Award partial credit per author: 0.10 per author
    try:
        authors_with_xe = 0
        for author in REQUIRED_AUTHORS:
            if xe_counter.get(author, 0) > 0:
                authors_with_xe += 1
                print(f"PASS: XE entry found for '{author}' (count: {xe_counter[author]})")
            else:
                print(f"FAIL: No XE entry found for '{author}'")

        if authors_with_xe == len(REQUIRED_AUTHORS):
            comp1_score = 0.50
            print(f"PASS: Component 1 — All 5 authors have XE entries ({comp1_score} pts)")
        elif authors_with_xe > 0:
            comp1_score = round(0.10 * authors_with_xe, 2)
            print(f"PARTIAL: Component 1 — {authors_with_xe}/5 authors have XE entries ({comp1_score} pts)")
        else:
            comp1_score = 0.0
            print(f"FAIL: Component 1 — No authors have XE entries")
        total_score += comp1_score
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: INDEX field is present in the document (0.20 points)
    try:
        has_index_field = bool(re.search(r'INDEX', body_xml))
        if has_index_field:
            print(f"PASS: Component 2 — INDEX field found in document (0.20 pts)")
            total_score += 0.20
        else:
            print(f"FAIL: Component 2 — No INDEX field found in document")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Author Index heading/section exists at end of document (0.15 points)
    # Check that there is a heading with "Index" or "Author Index" near the end
    try:
        last_paras = doc.paragraphs[-5:] if len(doc.paragraphs) >= 5 else doc.paragraphs
        has_index_heading = False
        for p in last_paras:
            if p.style and 'Heading' in p.style.name and 'index' in p.text.lower():
                has_index_heading = True
                break
            # Also check for non-heading paragraph that says "Author Index" as a title
            if 'author index' in p.text.lower() and len(p.text.strip()) < 50:
                has_index_heading = True
                break

        if has_index_heading:
            print(f"PASS: Component 3 — Author Index heading found at end of document (0.15 pts)")
            total_score += 0.15
        else:
            # Broader check: any paragraph near the end with "index" in it
            # that is a heading style
            all_headings = [(i, p) for i, p in enumerate(doc.paragraphs)
                          if p.style and 'Heading' in p.style.name and 'index' in p.text.lower()]
            if all_headings:
                print(f"PARTIAL: Component 3 — Index heading found but not at very end (paragraph {all_headings[-1][0]}) (0.10 pts)")
                total_score += 0.10
            else:
                print(f"FAIL: Component 3 — No Author Index heading found")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Sufficient XE entry coverage (0.15 points)
    # The 5 authors should have multiple XE entries reflecting their multiple mentions
    # At least 10 total XE entries across all 5 authors (task has many citations)
    try:
        total_xe = sum(xe_counter.get(a, 0) for a in REQUIRED_AUTHORS)
        if total_xe >= 10 and len(found_authors) == 5:
            print(f"PASS: Component 4 — {total_xe} total XE entries across all 5 authors (0.15 pts)")
            total_score += 0.15
        elif total_xe >= 5 and len(found_authors) >= 3:
            print(f"PARTIAL: Component 4 — {total_xe} XE entries, {len(found_authors)} authors covered (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 4 — Only {total_xe} total XE entries across {len(found_authors)} authors")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state("libreoffice_writer")

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
