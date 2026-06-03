"""
FINAL REWARD SCRIPT - SUCCESS
Task: In LibreOffice Writer I discovered that every italicized word in paragraphs 2, 3, and 4 is stuck at 10 pt. Everywhere else I’m using 13 pt. How can I quickly change ONLY the italic text in those three paragraphs to exactly 13 pt without clicking each word manually?
Generated: 2025-09-10 14:45:57
Status: success
Model: azure-o3
Total Steps: 4
"""

import os
from docx import Document
from docx.shared import Pt,Length

"""
Reward Script for LibreOffice Writer Italic Font-Size Fix Task
Task Recap:
The user needed to change ONLY the italicised words that appear in paragraphs 2, 3 and 4
from 10 pt to 13 pt, leaving every other part of the document unchanged.
This script awards a progressive score (0.0-1.0) based on how accurately
those specific italic runs were updated and whether no other italic text
is left at the wrong size.

Scoring Breakdown (adds up, then capped at 1.0):
  • 0.30 pts – Italic runs actually present in paragraphs 2-4 (prerequisite to judge change)
  • 0.40 pts – Proportion of italic runs in paragraphs 2-4 that now have a ≥12.6 pt size
               (i.e. effectively 13 pt).  full 0.40 when all are correct.
  • 0.30 pts – Quality check across the whole file: no italic run anywhere
               is still smaller than 12.6 pt.  Partial credit if only some are wrong.

The script uses python-docx to examine each run, checks its italic flag and
its explicit font size (run.font.size).  The size value in a DOCX run is
stored as a Length (EMU).  13 pt equals 165100 EMU.  Anything below
~160000 EMU (~12.6 pt) is considered still “too small”.

The script prints detailed diagnostics and finally prints the reward as:
    REWARD: <score>
exactly matching the evaluation harness expectations.
"""

# ---------------------------------------------------------------------------
# Helper Constants
# ---------------------------------------------------------------------------
# 13 pt in EMUs = 13 * 12700
EMUS_13PT = int(round(13 * 12700))  # 165100
# Threshold a bit lower to allow rounding differences (≈12.6 pt)
MIN_ACCEPTABLE_EMUS = int(round(12.6 * 12700))  # 160020


def _emu_from_length(length_obj):
    """Convert python-docx Length or None to an integer EMU value (or None)."""
    if length_obj is None:
        return None
    # python-docx Length behaves like int but provides .pt property as float pts.
    if isinstance(length_obj, Length):
        return int(length_obj)  # already EMU in int
    # Fallback: if object has .pt, convert via points
    if hasattr(length_obj, "pt"):
        return int(round(length_obj.pt * 12700))
    # If it is already int/float assume EMU
    try:
        return int(length_obj)
    except Exception:
        return None


def verify_writer_italic_fix(file_path: str) -> float:
    """Main verification logic – returns a float reward between 0.0 and 1.0."""
    total_score = 0.0
    max_score = 1.0

    # -------------------------------------------------------------------
    # Prerequisite: file must exist & load – NO POINTS AWARDED for this
    # -------------------------------------------------------------------
    if not os.path.exists(file_path):
        print(f"✗ File not found: {file_path}")
        return 0.0

    try:
        doc = Document(file_path)
    except Exception as exc:
        print(f"✗ Unable to open DOCX: {exc}")
        return 0.0

    print(f"✓ Loaded document – {len(doc.paragraphs)} paragraphs detected")

    # -------------------------------------------------------------------
    # Analyse every run
    # -------------------------------------------------------------------
    italic_runs_p234 = 0          # how many italic runs in paragraphs 2-4 total
    italic_runs_p234_correct = 0  # of those, how many are correct size now

    total_italic_runs = 0         # italic runs in whole doc
    total_small_italic_runs = 0   # italic runs still smaller than threshold

    for p_idx, para in enumerate(doc.paragraphs, start=1):
        for run in para.runs:
            if not run.italic:
                continue  # only interested in italic text

            total_italic_runs += 1

            size_emu = _emu_from_length(run.font.size)
            size_info = "(inherit)" if size_emu is None else f"{size_emu} EMU"

            # Global tiny-size check (all paragraphs)
            if size_emu is not None and size_emu < MIN_ACCEPTABLE_EMUS:
                total_small_italic_runs += 1
                print(f"  ✗ Small italic run in paragraph {p_idx} »{run.text}« {size_info}")

            # Specific paragraphs 2-4 analysis
            if p_idx in (2, 3, 4):
                italic_runs_p234 += 1
                # Treat size==None as acceptable (font size inherited from 13-pt style)
                if size_emu is None or size_emu >= MIN_ACCEPTABLE_EMUS:
                    italic_runs_p234_correct += 1
                else:
                    print(f"    ✗ Paragraph {p_idx} italic run still too small: »{run.text}« {size_info}")

    # -------------------------------------------------------------------
    # Scoring Section
    # -------------------------------------------------------------------
    print("\n--- Scoring ---")

    # (A) Presence of italic runs in paragraphs 2-4 (0.30)
    if italic_runs_p234 > 0:
        total_score += 0.30
        print(f"✓ Italic runs detected in paragraphs 2-4 ({italic_runs_p234}) → +0.30")
    else:
        print("✗ No italic runs found in paragraphs 2-4 → +0.00")

    # (B) Correctly sized italic runs in 2-4 (up to 0.40)
    if italic_runs_p234 > 0:
        ratio_correct = italic_runs_p234_correct / italic_runs_p234
        pts = 0.40 * ratio_correct
        total_score += pts
        print(f"✓ Size correctness in paragraphs 2-4: {italic_runs_p234_correct}/{italic_runs_p234} → +{pts:.2f}")
    else:
        print("(skipped) could not evaluate correctness without italic runs → +0.00")

    # (C) No small italic runs anywhere else (up to 0.30)
    if total_italic_runs == 0:
        # Unusual case: no italic text at all – give neutral credit
        print("(neutral) Document contains no italic text → +0.15")
        total_score += 0.15  # half credit – nothing to be wrong
    else:
        ok_ratio = 1 - (total_small_italic_runs / total_italic_runs)
        pts = 0.30 * ok_ratio
        total_score += pts
        if total_small_italic_runs == 0:
            print(f"✓ All italic runs acceptable size ({total_italic_runs} runs) → +0.30")
        else:
            print(f"⚠ {total_small_italic_runs} / {total_italic_runs} italic runs still small → +{pts:.2f}")

    # Cap final score at 1.0
    final_score = min(max_score, round(total_score, 4))
    print(f"\nTotal score: {final_score}")
    return final_score


# ---------------------------------------------------------------------------
# MAIN EXECUTION (called automatically when script runs)
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    # Expected file path (provided in the task context)
    DEFAULT_PATH = "/home/user/in_libreoffice_writer_i_discovered_that_every_italicized_word_in_paragraphs_2_3_and_4_is_stuck_at_10.docx"

    target_path = os.environ.get("WRITER_TASK_FILE", DEFAULT_PATH)
    reward_value = verify_writer_italic_fix(target_path)
    print(f"REWARD: {reward_value}")

