"""
Initial Setup: Accept formatting-only tracked changes in Brand Guidelines document
Task ID: writer_rm_017
Domain: libreoffice_writer

Creates a Brand_Guidelines.docx with 15 tracked changes:
  - 6 formatting-only changes (font size, bold, color, italic, underline adjustments)
  - 9 text content changes (insertions, deletions)
"""

import os
import shlex
import subprocess
import time
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_rm_017'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def qn(tag):
    prefix, local = tag.split(':')
    ns = {'w': W}
    return f'{{{ns[prefix]}}}{local}'

def launch_gui(command: str, delay_sec: float = 1.0):
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def build_base_document():
    """Build the base Brand Guidelines document with realistic content."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Nexus Corp Brand Guidelines', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 1: Brand Overview
    doc.add_heading('1. Brand Overview', level=1)
    p = doc.add_paragraph()
    p.add_run('Nexus Corp is a global technology company dedicated to innovation and excellence. ').font.size = Pt(11)
    p.add_run('Founded in 2015, our mission is to bridge the gap between cutting-edge technology and everyday solutions.').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Our brand identity reflects professionalism, reliability, and forward-thinking design principles that guide all our visual and verbal communications.').font.size = Pt(11)

    # Section 2: Logo Usage
    doc.add_heading('2. Logo Usage', level=1)
    p = doc.add_paragraph()
    p.add_run('The Nexus Corp logo must always be displayed with adequate clear space around it. ').font.size = Pt(11)
    p.add_run('The minimum clear space is defined as the height of the letter "N" in the wordmark, measured on all four sides.').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Do not stretch, rotate, or alter the proportions of the logo under any circumstances. ').font.size = Pt(11)
    p.add_run('Approved logo files are available in the shared brand assets folder.').font.size = Pt(11)

    # Section 3: Color Palette
    doc.add_heading('3. Color Palette', level=1)
    p = doc.add_paragraph()
    p.add_run('Primary Blue: #1A5276 (Pantone 302 C)').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Accent Gold: #D4AC0D (Pantone 7405 C)').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Neutral Gray: #5D6D7E (Pantone Cool Gray 10 C)').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('These colors must be used consistently across all marketing materials, presentations, and digital platforms.').font.size = Pt(11)

    # Section 4: Typography
    doc.add_heading('4. Typography', level=1)
    p = doc.add_paragraph()
    r = p.add_run('Primary Typeface: ')
    r.font.size = Pt(11)
    r = p.add_run('Montserrat')
    r.font.size = Pt(11)
    r.bold = True
    r = p.add_run(' for headings and display text.')
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run('Body Typeface: ')
    r.font.size = Pt(11)
    r = p.add_run('Open Sans')
    r.font.size = Pt(11)
    r.bold = True
    r = p.add_run(' for body copy and long-form reading.')
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Minimum font size for body text is 10pt. Heading sizes should follow the hierarchy: H1 at 24pt, H2 at 18pt, H3 at 14pt.').font.size = Pt(11)

    # Section 5: Voice and Tone
    doc.add_heading('5. Voice and Tone', level=1)
    p = doc.add_paragraph()
    p.add_run('Our brand voice is confident, approachable, and knowledgeable. ').font.size = Pt(11)
    p.add_run('We speak directly to our audience using clear, jargon-free language.').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Avoid overly technical terms when addressing general audiences. ').font.size = Pt(11)
    p.add_run('Use active voice and present tense whenever possible to maintain energy and directness.').font.size = Pt(11)

    # Section 6: Digital Standards
    doc.add_heading('6. Digital Standards', level=1)
    p = doc.add_paragraph()
    p.add_run('All web content must follow WCAG 2.1 AA accessibility standards. ').font.size = Pt(11)
    p.add_run('Color contrast ratios must meet or exceed 4.5:1 for normal text and 3:1 for large text.').font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run('Social media templates are available in the brand portal under the Digital Assets section.').font.size = Pt(11)

    return doc


def add_rpr_change(run_elem, rev_id, author, date):
    """Add a rPrChange to a run, recording its CURRENT rPr as the old state.
    Returns the rPrChange element so caller can set old formatting."""
    rpr = run_elem.find(qn('w:rPr'))
    if rpr is None:
        rpr = etree.SubElement(run_elem, qn('w:rPr'))
        run_elem.insert(0, rpr)
    rpr_change = etree.SubElement(rpr, qn('w:rPrChange'))
    rpr_change.set(qn('w:id'), str(rev_id))
    rpr_change.set(qn('w:author'), author)
    rpr_change.set(qn('w:date'), date)
    old_rpr = etree.SubElement(rpr_change, qn('w:rPr'))
    return rpr, old_rpr


def make_ins(rev_id, author, date, text):
    """Create a w:ins element with a run containing text."""
    ins = etree.Element(qn('w:ins'))
    ins.set(qn('w:id'), str(rev_id))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date)
    run = etree.SubElement(ins, qn('w:r'))
    t = etree.SubElement(run, qn('w:t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return ins


def make_del(rev_id, author, date, text):
    """Create a w:del element with a run containing delText."""
    d = etree.Element(qn('w:del'))
    d.set(qn('w:id'), str(rev_id))
    d.set(qn('w:author'), author)
    d.set(qn('w:date'), date)
    run = etree.SubElement(d, qn('w:r'))
    t = etree.SubElement(run, qn('w:delText'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return d


def create_initial():
    doc = build_base_document()
    doc.save(OUTPUT)

    # Reopen to manipulate XML
    doc = Document(OUTPUT)
    body = doc.element.body
    paras = body.findall(qn('w:p'))

    AUTHOR = 'Elena Rodriguez'
    DATE = '2025-11-14T09:30:00Z'
    rid = 100

    # Map paragraph indices (verified):
    # 0=Title, 1=H1(Brand Overview), 2=p_overview1, 3=p_overview2,
    # 4=H1(Logo), 5=p_logo1, 6=p_logo2,
    # 7=H1(Color), 8=p_blue, 9=p_gold, 10=p_gray, 11=p_coloruse,
    # 12=H1(Typography), 13=p_typeface1, 14=p_typeface2, 15=p_fontsizes,
    # 16=H1(Voice), 17=p_voice1, 18=p_voice2,
    # 19=H1(Digital), 20=p_wcag, 21=p_social

    # ===== 6 FORMATTING-ONLY TRACKED CHANGES =====

    # FC1: para[2] run[0] - Make first sentence of Brand Overview bold
    # New: bold=True. Old: no bold.
    r = paras[2].findall(qn('w:r'))[0]
    rpr, old_rpr = add_rpr_change(r, rid, AUTHOR, DATE); rid += 1
    etree.SubElement(rpr, qn('w:b'))

    # FC2: para[5] run[1] - Change logo clear space sentence font size 11pt->12pt
    # New: sz=24. Old: sz=22.
    r = paras[5].findall(qn('w:r'))[1]
    rpr, old_rpr = add_rpr_change(r, rid, AUTHOR, DATE); rid += 1
    sz = etree.SubElement(rpr, qn('w:sz'))
    sz.set(qn('w:val'), '24')
    szCs = etree.SubElement(rpr, qn('w:szCs'))
    szCs.set(qn('w:val'), '24')
    old_sz = etree.SubElement(old_rpr, qn('w:sz'))
    old_sz.set(qn('w:val'), '22')
    old_szCs = etree.SubElement(old_rpr, qn('w:szCs'))
    old_szCs.set(qn('w:val'), '22')

    # FC3: para[8] run[0] - Color the "Primary Blue" text blue (#1A5276)
    # New: color=1A5276. Old: no color.
    r = paras[8].findall(qn('w:r'))[0]
    rpr, old_rpr = add_rpr_change(r, rid, AUTHOR, DATE); rid += 1
    c = etree.SubElement(rpr, qn('w:color'))
    c.set(qn('w:val'), '1A5276')

    # FC4: para[13] run[1] - Make "Montserrat" italic (it's already bold)
    # New: italic. Old: just bold.
    r = paras[13].findall(qn('w:r'))[1]
    rpr, old_rpr = add_rpr_change(r, rid, AUTHOR, DATE); rid += 1
    etree.SubElement(rpr, qn('w:i'))
    # Record old state had bold but no italic
    etree.SubElement(old_rpr, qn('w:b'))

    # FC5: para[9] run[0] - Make "Accent Gold" text bold
    # New: bold. Old: no bold.
    r = paras[9].findall(qn('w:r'))[0]
    rpr, old_rpr = add_rpr_change(r, rid, AUTHOR, DATE); rid += 1
    etree.SubElement(rpr, qn('w:b'))

    # FC6: para[20] run[1] - Underline contrast ratios sentence
    # New: underline. Old: no underline.
    r = paras[20].findall(qn('w:r'))[1]
    rpr, old_rpr = add_rpr_change(r, rid, AUTHOR, DATE); rid += 1
    u = etree.SubElement(rpr, qn('w:u'))
    u.set(qn('w:val'), 'single')

    # ===== 9 TEXT CONTENT TRACKED CHANGES =====

    # TC1: para[2] - Insert "and sustainability " after first run
    ins = make_ins(rid, AUTHOR, DATE, 'and sustainability '); rid += 1
    runs_p2 = paras[2].findall(qn('w:r'))
    runs_p2[0].addnext(ins)

    # TC2: para[2] - Delete "everyday " (marks existing text as deleted)
    d = make_del(rid, AUTHOR, DATE, 'everyday '); rid += 1
    paras[2].append(d)

    # TC3: para[5] - Insert "mandatory " at end
    ins = make_ins(rid, AUTHOR, DATE, 'mandatory '); rid += 1
    paras[5].append(ins)

    # TC4: para[6] - Delete "under any circumstances"
    d = make_del(rid, AUTHOR, DATE, 'under any circumstances'); rid += 1
    paras[6].append(d)

    # TC5: para[11] - Insert "or digital displays " at end
    ins = make_ins(rid, AUTHOR, DATE, 'or digital displays '); rid += 1
    paras[11].append(ins)

    # TC6: para[15] - Insert "recommended " at end
    ins = make_ins(rid, AUTHOR, DATE, 'recommended '); rid += 1
    paras[15].append(ins)

    # TC7: para[17] - Delete "jargon-free "
    d = make_del(rid, AUTHOR, DATE, 'jargon-free '); rid += 1
    paras[17].append(d)

    # TC8: para[18] - Insert "and concise " at end
    ins = make_ins(rid, AUTHOR, DATE, 'and concise '); rid += 1
    paras[18].append(ins)

    # TC9: para[20] - Insert "strictly " at end
    ins = make_ins(rid, AUTHOR, DATE, 'strictly '); rid += 1
    paras[20].append(ins)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')
    print(f'Total tracked changes: 15 (6 formatting-only, 9 text content)')

    # Verify
    doc2 = Document(OUTPUT)
    body2 = doc2.element.body
    fmt_changes = len(body2.findall(f'.//{qn("w:rPrChange")}'))
    ins_changes = len(body2.findall(f'.//{qn("w:ins")}'))
    del_changes = len(body2.findall(f'.//{qn("w:del")}'))
    print(f'Verification - Format changes: {fmt_changes}, Insertions: {ins_changes}, Deletions: {del_changes}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
