"""
Initial Setup: Insert a cross-reference to Bookmark 'sec_results'
Task ID: writer_tm_072
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_tm_072'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_bookmark(paragraph, bookmark_name, text):
    """Add a bookmark wrapping text in a paragraph."""
    import random
    bm_id = str(random.randint(1000, 99999))
    run = paragraph.add_run(text)

    # Insert bookmarkStart before the run
    bm_start = paragraph._element.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): bm_id,
        qn('w:name'): bookmark_name,
    })
    run._element.addprevious(bm_start)

    # Insert bookmarkEnd after the run
    bm_end = paragraph._element.makeelement(qn('w:bookmarkEnd'), {
        qn('w:id'): bm_id,
    })
    run._element.addnext(bm_end)

    return run


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # === PAGE 1: Title and Abstract ===
    title = doc.add_heading('Evaluating the Impact of Remote Work on Team Productivity in Software Engineering Organizations', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = authors.add_run('Dr. Elena Vasquez, Prof. James Whitfield, Dr. Anika Sharma')
    run.font.size = Pt(11)

    affil = doc.add_paragraph()
    affil.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = affil.add_run('Department of Computer Science, Pacific Northwest University\nPortland, OR 97201, USA')
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    abs_heading = doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This study investigates the relationship between remote work arrangements and team '
        'productivity metrics in software engineering organizations. Drawing on a longitudinal '
        'dataset of 847 software development teams across 32 companies, we examine how the '
        'transition to remote and hybrid work models influenced sprint velocity, code review '
        'turnaround times, and bug resolution rates over a 24-month period from January 2023 '
        'to December 2024. Our findings suggest a nuanced picture where initial productivity '
        'dips were followed by sustained improvements in teams that adopted structured '
        'communication protocols.'
    )
    doc.add_paragraph(
        'Keywords: remote work, software engineering, team productivity, agile development, '
        'distributed teams, organizational behavior'
    ).runs[0].font.italic = True

    add_page_break(doc)

    # === PAGE 2: Introduction ===
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The global shift toward remote work, accelerated by the COVID-19 pandemic, has '
        'fundamentally transformed how software engineering teams collaborate and deliver '
        'products. Prior to 2020, only 16% of technology companies offered fully remote '
        'positions (Bureau of Labor Statistics, 2019). By mid-2023, this figure had risen '
        'to 62%, with an additional 24% offering hybrid arrangements (McKinsey Global '
        'Institute, 2023).'
    )
    doc.add_paragraph(
        'Despite the widespread adoption of remote work, there remains considerable debate '
        'about its long-term impact on team productivity. Proponents argue that remote work '
        'eliminates commuting overhead, reduces interruptions, and provides developers with '
        'deeper focus time (Newport, 2022; Bloom et al., 2023). Critics counter that the '
        'loss of spontaneous face-to-face interaction diminishes creative problem-solving '
        'and slows knowledge transfer within teams (Waber et al., 2021).'
    )
    doc.add_paragraph(
        'Several large-scale studies have attempted to measure these effects, but most '
        'rely on self-reported productivity measures or single-company datasets, limiting '
        'generalizability. Our study addresses this gap by leveraging objective metrics '
        'derived from project management tools (Jira, Linear, and Shortcut) across a diverse '
        'sample of organizations ranging from 50-person startups to Fortune 500 enterprises.'
    )

    # This is the key paragraph where cursor should be - on page 2
    p_ref = doc.add_paragraph(
        'Our findings are discussed in '
    )
    # Cursor is here, after "Our findings are discussed in " - no cross-reference yet

    doc.add_paragraph(
        'The remainder of this paper is organized as follows. Section 2 reviews the relevant '
        'literature. Section 3 describes our methodology and data collection procedures. '
        'Section 4 presents detailed results. Section 5 discusses implications for both '
        'researchers and practitioners.'
    )

    add_page_break(doc)

    # === PAGE 3: Literature Review ===
    doc.add_heading('2. Literature Review', level=1)
    doc.add_heading('2.1 Historical Context of Remote Work in Technology', level=2)
    doc.add_paragraph(
        'The concept of telecommuting in the technology sector predates the pandemic by '
        'several decades. Nilles (1975) first proposed the idea of "telecommuting" as a '
        'strategy to reduce urban congestion and energy consumption. However, widespread '
        'adoption remained limited due to technological constraints and managerial resistance '
        '(Bailey & Kurland, 2002).'
    )
    doc.add_paragraph(
        'The emergence of cloud-based collaboration tools in the 2010s, including Slack, '
        'GitHub, and Zoom, gradually reduced the barriers to distributed work. Companies '
        'such as GitLab, Automattic, and Basecamp demonstrated that fully remote software '
        'development was not only feasible but could offer competitive advantages in talent '
        'acquisition and employee retention (Fried & Hansson, 2013).'
    )

    doc.add_heading('2.2 Productivity Measurement in Software Engineering', level=2)
    doc.add_paragraph(
        'Measuring developer productivity has been a persistent challenge in software '
        'engineering research. Traditional metrics such as lines of code (LOC) have been '
        'widely criticized for incentivizing code bloat rather than quality (DeMarco & '
        'Lister, 2013). More recent approaches use composite metrics that balance throughput '
        '(story points delivered, pull requests merged) with quality indicators (defect '
        'escape rate, code review coverage).'
    )
    doc.add_paragraph(
        'The DORA metrics framework (Forsgren et al., 2018) has gained significant traction, '
        'measuring deployment frequency, lead time for changes, change failure rate, and mean '
        'time to recovery. Our study incorporates elements of DORA alongside additional '
        'team-level metrics specific to the remote work context.'
    )

    add_page_break(doc)

    # === PAGE 4: Methodology ===
    doc.add_heading('3. Methodology', level=1)
    doc.add_heading('3.1 Research Design', level=2)
    doc.add_paragraph(
        'We employed a mixed-methods longitudinal research design combining quantitative '
        'analysis of project management data with qualitative interviews of engineering '
        'managers. The quantitative component draws on automated data extraction from Jira, '
        'Linear, and Shortcut instances, supplemented by Git repository analytics.'
    )
    doc.add_paragraph(
        'Participating organizations were recruited through industry partnerships and '
        'professional networks between March and June 2023. Eligibility criteria included: '
        '(a) employing at least 20 software developers, (b) using an agile project management '
        'tool with at least 12 months of historical data, and (c) having undergone a transition '
        'from primarily in-office to remote or hybrid work during 2020-2022.'
    )

    doc.add_heading('3.2 Data Collection', level=2)
    doc.add_paragraph(
        'Quantitative data were collected via API integrations with participating organizations\' '
        'project management platforms. We extracted sprint-level metrics including velocity '
        '(story points completed per sprint), cycle time (days from ticket creation to '
        'completion), code review turnaround (hours from PR submission to first review), '
        'and bug resolution rate (percentage of bugs resolved within SLA). All data were '
        'anonymized at the point of collection using a secure hashing protocol approved '
        'by our institutional review board (IRB Protocol #2023-0847).'
    )

    doc.add_heading('3.3 Sample Characteristics', level=2)
    doc.add_paragraph(
        'The final sample comprised 847 teams from 32 organizations. Team sizes ranged from '
        '4 to 18 members (M = 7.3, SD = 2.8). Organizations spanned multiple sub-sectors: '
        'enterprise SaaS (38%), consumer applications (22%), fintech (19%), healthcare '
        'technology (12%), and infrastructure/DevOps tools (9%). Geographic distribution '
        'included North America (56%), Europe (28%), and Asia-Pacific (16%).'
    )

    add_page_break(doc)

    # === PAGE 5: Methodology continued / Data Analysis ===
    doc.add_heading('3.4 Analytical Framework', level=2)
    doc.add_paragraph(
        'We employed a difference-in-differences (DiD) estimation strategy to isolate the '
        'causal effect of remote work adoption on productivity metrics. The treatment group '
        'consisted of teams that transitioned to fully remote work, while the control group '
        'comprised teams that returned to primarily in-office arrangements after initial '
        'lockdowns. Hybrid teams were analyzed as a separate treatment arm.'
    )
    doc.add_paragraph(
        'The baseline model specification is as follows: Y_it = alpha + beta * Remote_it + '
        'gamma * X_it + delta_i + tau_t + epsilon_it, where Y_it represents the productivity '
        'metric for team i in period t, Remote_it is the treatment indicator, X_it is a '
        'vector of time-varying covariates (team size, project complexity, sprint duration), '
        'delta_i captures team fixed effects, and tau_t captures time fixed effects.'
    )
    doc.add_paragraph(
        'To address potential selection bias, we employed propensity score matching (PSM) '
        'as a robustness check. Teams were matched on observable pre-treatment characteristics '
        'including historical velocity, team tenure, technology stack, and organizational size. '
        'The propensity scores were estimated using a logistic regression model with a '
        'caliper of 0.02 standard deviations.'
    )

    doc.add_heading('3.5 Qualitative Component', level=2)
    doc.add_paragraph(
        'Semi-structured interviews were conducted with 48 engineering managers across 24 '
        'organizations. Interviews lasted 45-60 minutes and covered topics including team '
        'communication patterns, onboarding practices, and perceived productivity changes. '
        'Transcripts were analyzed using thematic coding with NVivo software. Inter-rater '
        'reliability was established with two independent coders achieving a Cohen\'s kappa '
        'of 0.83.'
    )

    add_page_break(doc)

    # === PAGE 6: Results and Discussion ===
    results_heading = doc.add_heading('4. Results and Discussion', level=1)

    # Add bookmark 'sec_results' to this heading
    bm_id = '10001'
    bm_start = results_heading._element.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): bm_id,
        qn('w:name'): 'sec_results',
    })
    bm_end = results_heading._element.makeelement(qn('w:bookmarkEnd'), {
        qn('w:id'): bm_id,
    })
    # Insert at beginning and end of heading element
    first_run = results_heading._element.findall(qn('w:r'))
    if first_run:
        first_run[0].addprevious(bm_start)
        first_run[-1].addnext(bm_end)
    else:
        results_heading._element.append(bm_start)
        results_heading._element.append(bm_end)

    doc.add_heading('4.1 Overall Productivity Trends', level=2)
    doc.add_paragraph(
        'Figure 1 presents the aggregate sprint velocity trends across all three work '
        'arrangement groups. Fully remote teams experienced an initial productivity decline '
        'of 12.4% (95% CI: 8.7%-16.1%) in the first quarter following transition, consistent '
        'with prior findings (Yang et al., 2022). However, by the third quarter, remote '
        'teams had recovered to pre-transition levels, and by the sixth quarter, they '
        'demonstrated a sustained 7.2% improvement (95% CI: 3.8%-10.6%, p < 0.001).'
    )
    doc.add_paragraph(
        'Hybrid teams showed a more modest initial decline of 5.8% (95% CI: 2.1%-9.5%) '
        'and achieved a 4.1% long-term improvement (95% CI: 1.2%-7.0%, p = 0.006). In-office '
        'teams showed no statistically significant change in velocity over the same period '
        '(beta = 0.3%, p = 0.847), serving as an effective control group.'
    )

    doc.add_heading('4.2 Code Review Turnaround', level=2)
    doc.add_paragraph(
        'Code review turnaround times exhibited a different pattern. Remote teams experienced '
        'a persistent increase in median review time from 4.2 hours to 6.8 hours (a 62% '
        'increase, p < 0.001). This effect was partially mitigated in teams that adopted '
        'asynchronous code review protocols with structured review windows, where the '
        'increase was limited to 23% (median 5.2 hours).'
    )

    doc.add_heading('4.3 Bug Resolution Rates', level=2)
    doc.add_paragraph(
        'Bug resolution rates within SLA improved by 9.3% for remote teams (p = 0.003), '
        'which we attribute to reduced context-switching enabled by fewer in-person '
        'interruptions. This finding aligns with the deep work hypothesis proposed by '
        'Newport (2022) and suggests that certain categories of focused engineering work '
        'benefit from the uninterrupted time afforded by remote arrangements.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready: open in LibreOffice Writer
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
