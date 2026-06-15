"""
Reward Script: Accept three tracked changes (text corrections), reject two tracked changes (budget figures),
               and leave remaining 5 tracked changes unresolved.
Task ID: writer_struct_070
Domain: libreoffice_writer
Scoring:
  Component 1: Three accepted changes are finalized (0.4 pts)
               - 'Fiscal Year 2025', 'division head', 'quarterly review' are in plain text
               - No tracked insertion/deletion markers for these changes
  Component 2: Two rejected changes are reverted (0.4 pts)
               - '$2.5M' and '$500K' are present as plain text (original values restored)
               - '$3.0M' and '$750K' are NOT present as tracked insertions
  Component 3: Remaining 5 tracked changes are still unresolved (0.2 pts)
               - Exactly 5 insertion/deletion pairs remain in the document
"""

import os
from docx import Document

WORKDIR = '/home/user'
FILE_PATH = '/home/user/Desktop/annual_budget_review.docx'

# Namespaces for XML traversal
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': W_NS}


def get_tracked_insertions(root):
    """Return list of (id, text) tuples for all tracked insertions in the document."""
    results = []
    for ins in root.findall('.//w:ins', NS):
        texts = [t.text for t in ins.findall('.//w:t', NS) if t.text]
        ins_id = ins.get(f'{{{W_NS}}}id', 'no-id')
        combined = ''.join(texts)
        results.append((ins_id, combined))
    return results


def get_tracked_deletions(root):
    """Return list of (id, text) tuples for all tracked deletions in the document."""
    results = []
    for d in root.findall('.//w:del', NS):
        texts = [t.text for t in d.findall('.//w:delText', NS) if t.text]
        del_id = d.get(f'{{{W_NS}}}id', 'no-id')
        combined = ''.join(texts)
        results.append((del_id, combined))
    return results


def get_plain_text_paragraphs(doc):
    """Return concatenated plain text from all paragraphs, including tracked insertions
    but excluding tracked deletions (i.e., what the document 'looks like' after accepting all changes)."""
    # doc.paragraphs[i].text returns the visible text with accepted changes applied
    return '\n'.join(p.text for p in doc.paragraphs)


def verify_task():
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    # --- Load the document ---
    if not os.path.exists(FILE_PATH):
        print(f"CRITICAL: File not found: {FILE_PATH}")
        print("REWARD: 0.0")
        return 0.0

    try:
        doc = Document(FILE_PATH)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {FILE_PATH}: {e}")
        print("REWARD: 0.0")
        return 0.0

    root = doc.element
    insertions = get_tracked_insertions(root)
    deletions = get_tracked_deletions(root)

    # Build sets of insertion texts and deletion texts for easy lookup
    insertion_texts = set(text for _, text in insertions)
    deletion_texts = set(text for _, text in deletions)
    full_doc_text = get_plain_text_paragraphs(doc)

    print(f"INFO: Found {len(insertions)} tracked insertions, {len(deletions)} tracked deletions")
    print(f"INFO: Insertion texts: {insertion_texts}")
    print(f"INFO: Deletion texts: {deletion_texts}")

    # ----------------------------------------------------------------
    # Component 1: Three accepted changes are finalized (0.4 points)
    # Task requires accepting:
    #   - 'Fiscal Year 2024' → 'Fiscal Year 2025'  (page 1)
    #   - 'department head'  → 'division head'      (page 3)
    #   - 'annual review'    → 'quarterly review'   (page 5)
    # After accepting: the new text is in plain doc, old text NOT in deletions,
    # new text NOT in insertions (it's been finalized).
    # ----------------------------------------------------------------
    try:
        accepted_ok = 0

        # Check 'Fiscal Year 2025' is finalized:
        #   - present in document text (visible)
        #   - NOT still a tracked insertion (meaning it was finalized)
        #   - 'Fiscal Year 2024' NOT still a tracked deletion
        fy2025_in_text = 'Fiscal Year 2025' in full_doc_text
        fy2025_still_ins = 'Fiscal Year 2025' in insertion_texts
        fy2024_still_del = 'Fiscal Year 2024' in deletion_texts
        fy_accepted = fy2025_in_text and not fy2025_still_ins and not fy2024_still_del

        if fy_accepted:
            print("PASS: 'Fiscal Year 2025' accepted (finalized, no longer tracked)")
            accepted_ok += 1
        else:
            print(f"FAIL: 'Fiscal Year 2025' NOT finalized. "
                  f"in_text={fy2025_in_text}, still_tracked_ins={fy2025_still_ins}, "
                  f"old_still_del={fy2024_still_del}")

        # Check 'division head' is finalized:
        divhead_in_text = 'division head' in full_doc_text
        divhead_still_ins = 'division head' in insertion_texts
        dephead_still_del = 'department head' in deletion_texts
        divhead_accepted = divhead_in_text and not divhead_still_ins and not dephead_still_del

        if divhead_accepted:
            print("PASS: 'division head' accepted (finalized, no longer tracked)")
            accepted_ok += 1
        else:
            print(f"FAIL: 'division head' NOT finalized. "
                  f"in_text={divhead_in_text}, still_tracked_ins={divhead_still_ins}, "
                  f"old_still_del={dephead_still_del}")

        # Check 'quarterly review' is finalized:
        qreview_in_text = 'quarterly review' in full_doc_text
        qreview_still_ins = 'quarterly review' in insertion_texts
        areview_still_del = 'annual review' in deletion_texts
        qreview_accepted = qreview_in_text and not qreview_still_ins and not areview_still_del

        if qreview_accepted:
            print("PASS: 'quarterly review' accepted (finalized, no longer tracked)")
            accepted_ok += 1
        else:
            print(f"FAIL: 'quarterly review' NOT finalized. "
                  f"in_text={qreview_in_text}, still_tracked_ins={qreview_still_ins}, "
                  f"old_still_del={areview_still_del}")

        if accepted_ok == 3:
            print(f"PASS: Component 1 — All 3 acceptance changes finalized (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 1 — Only {accepted_ok}/3 acceptance changes finalized")

    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # ----------------------------------------------------------------
    # Component 2: Two rejected changes are reverted (0.4 points)
    # Task requires rejecting:
    #   - '$2.5M' → '$3.0M'  (page 2)  → reject: revert to '$2.5M'
    #   - '$500K' → '$750K'  (page 4)  → reject: revert to '$500K'
    # After rejecting: old text IS in document text, new text NOT in insertions,
    # old text NOT in deletions.
    # ----------------------------------------------------------------
    try:
        rejected_ok = 0

        # Check '$2.5M' → '$3.0M' was rejected (revert to '$2.5M'):
        usd25m_in_text = '$2.5M' in full_doc_text
        usd30m_still_ins = '$3.0M' in insertion_texts
        usd25m_still_del = '$2.5M' in deletion_texts
        usd25m_rejected = usd25m_in_text and not usd30m_still_ins and not usd25m_still_del

        if usd25m_rejected:
            print("PASS: '$3.0M→$2.5M' change rejected (original '$2.5M' restored, not tracked)")
            rejected_ok += 1
        else:
            print(f"FAIL: '$2.5M' NOT properly rejected. "
                  f"$2.5M_in_text={usd25m_in_text}, $3.0M_still_ins={usd30m_still_ins}, "
                  f"$2.5M_still_del={usd25m_still_del}")

        # Check '$500K' → '$750K' was rejected (revert to '$500K'):
        usd500k_in_text = '$500K' in full_doc_text
        usd750k_still_ins = '$750K' in insertion_texts
        usd500k_still_del = '$500K' in deletion_texts
        usd500k_rejected = usd500k_in_text and not usd750k_still_ins and not usd500k_still_del

        if usd500k_rejected:
            print("PASS: '$750K→$500K' change rejected (original '$500K' restored, not tracked)")
            rejected_ok += 1
        else:
            print(f"FAIL: '$500K' NOT properly rejected. "
                  f"$500K_in_text={usd500k_in_text}, $750K_still_ins={usd750k_still_ins}, "
                  f"$500K_still_del={usd500k_still_del}")

        if rejected_ok == 2:
            print(f"PASS: Component 2 — Both rejection changes reverted (0.4 pts)")
            total_score += 0.4
        else:
            print(f"FAIL: Component 2 — Only {rejected_ok}/2 rejection changes reverted")

    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # ----------------------------------------------------------------
    # Component 3: Exactly 5 tracked changes remain unresolved (0.2 points)
    # The remaining changes (not accepted, not rejected) are:
    #   - 'compliance officer' → 'risk director'
    #   - '$4.2M' → '$4.8M'
    #   - 'conservative estimate' → 'baseline forecast'
    #   - '$15.2M' → '$16.8M'
    #   - 'finance team' → 'treasury department'
    # Expect exactly 5 ins + 5 del tracked change markers remaining.
    # ----------------------------------------------------------------
    try:
        remaining_ins_count = len(insertions)
        remaining_del_count = len(deletions)

        # Verify the specific remaining changes are present
        expected_remaining_ins = {
            'risk director', '$4.8M', 'baseline forecast', '$16.8M', 'treasury department'
        }
        expected_remaining_del = {
            'compliance officer', '$4.2M', 'conservative estimate', '$15.2M', 'finance team'
        }

        ins_match = insertion_texts == expected_remaining_ins
        del_match = deletion_texts == expected_remaining_del
        count_ok = remaining_ins_count == 5 and remaining_del_count == 5

        if count_ok and ins_match and del_match:
            print(f"PASS: Component 3 — Exactly 5 tracked change pairs remain unresolved (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 3 — Remaining tracked changes mismatch. "
                  f"ins_count={remaining_ins_count} (expected 5), "
                  f"del_count={remaining_del_count} (expected 5), "
                  f"ins_match={ins_match}, del_match={del_match}")
            if not ins_match:
                print(f"  Expected insertions: {expected_remaining_ins}")
                print(f"  Actual insertions:   {insertion_texts}")
            if not del_match:
                print(f"  Expected deletions: {expected_remaining_del}")
                print(f"  Actual deletions:   {deletion_texts}")

    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score:.1f}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point: verify the canonical task file on the VM
if __name__ == '__main__':
    verify_task()
