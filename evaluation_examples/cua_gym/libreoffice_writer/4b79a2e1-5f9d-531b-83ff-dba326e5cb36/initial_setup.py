"""
Initial Setup: Legal agreement with 25 defined terms in Definitions section, no index entries marked.
Task ID: writer_legal_058
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_058'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

DEFINED_TERMS = [
    ("Affiliate", "any entity that directly or indirectly controls, is controlled by, or is under common control with a Party, where 'control' means the ownership of more than fifty percent (50%) of the voting securities of such entity."),
    ("Agreement", "this Master Services Agreement, including all Exhibits, Schedules, and amendments hereto."),
    ("Applicable Law", "all laws, statutes, regulations, ordinances, rules, orders, decrees, and governmental requirements applicable to a Party in connection with its performance under this Agreement."),
    ("Business Day", "any day other than a Saturday, Sunday, or a day on which commercial banks in New York, New York are authorized or required by law to close."),
    ("Change of Control", "any merger, consolidation, sale of all or substantially all assets, or any transaction or series of related transactions in which more than fifty percent (50%) of the voting power of a Party is transferred."),
    ("Claim", "any third-party claim, demand, suit, action, or proceeding, including any regulatory investigation or audit, arising out of or relating to this Agreement."),
    ("Confidential Information", "all non-public information disclosed by one Party to the other Party, whether orally, in writing, or by inspection, that is designated as confidential or that reasonably should be understood to be confidential."),
    ("Damages", "all losses, liabilities, damages, costs, and expenses, including reasonable attorneys' fees and court costs, arising from or related to any breach of this Agreement."),
    ("Deliverables", "all work product, materials, documents, software, and other items to be provided by the Service Provider to the Client pursuant to a Statement of Work."),
    ("Effective Date", "the date first written above in the preamble of this Agreement, upon which the rights and obligations of the Parties shall commence."),
    ("Fee Schedule", "the schedule of fees, rates, and payment terms attached hereto as Exhibit B, as may be amended from time to time by mutual written agreement of the Parties."),
    ("Force Majeure Event", "any event beyond the reasonable control of a Party, including acts of God, war, terrorism, pandemic, epidemic, fire, flood, earthquake, labor disputes, or governmental actions."),
    ("Governing Law", "the laws of the State of Delaware, without regard to its conflict of laws principles, which shall govern the interpretation and enforcement of this Agreement."),
    ("Indemnified Party", "the Party seeking indemnification under Article VII of this Agreement, whether as the Client or the Service Provider."),
    ("Intellectual Property", "all patents, trademarks, copyrights, trade secrets, know-how, inventions, designs, software, databases, and all other intellectual property rights, whether registered or unregistered."),
    ("Key Personnel", "those individuals identified in Exhibit C who are essential to the performance of services under this Agreement and whose replacement requires prior written consent of the Client."),
    ("Liability Cap", "the maximum aggregate liability of either Party under this Agreement, which shall not exceed the total fees paid or payable during the twelve (12) month period preceding the event giving rise to the claim."),
    ("Material Breach", "a breach of any representation, warranty, covenant, or obligation under this Agreement that, if not cured within thirty (30) days after written notice, would substantially deprive the non-breaching Party of the benefit of this Agreement."),
    ("Notice", "any communication required or permitted under this Agreement, which shall be in writing and delivered by hand, overnight courier, certified mail, or electronic mail to the addresses set forth in Section 12.1."),
    ("Party", "individually, either the Client or the Service Provider, and 'Parties' shall mean both the Client and the Service Provider collectively."),
    ("Permitted Subcontractor", "any third party approved in writing by the Client to perform a portion of the services under a Statement of Work, subject to the terms and conditions of Article V."),
    ("Service Level Agreement", "the performance standards, metrics, availability requirements, and remedies for non-compliance set forth in Exhibit D attached hereto."),
    ("Statement of Work", "a document executed by both Parties that describes the specific services to be performed, the Deliverables, timeline, and fees for a particular engagement under this Agreement."),
    ("Term", "the initial period of this Agreement commencing on the Effective Date and continuing for three (3) years, unless earlier terminated in accordance with Article IX."),
    ("Termination for Convenience", "the right of either Party to terminate this Agreement or any Statement of Work without cause upon ninety (90) days' prior written Notice to the other Party."),
]


def create_initial():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('MASTER SERVICES AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Preamble
    preamble = doc.add_paragraph()
    preamble.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = preamble.add_run(
        'This Master Services Agreement (the "Agreement") is entered into as of March 15, 2025 '
        '(the "Effective Date"), by and between Meridian Technology Solutions, Inc., a Delaware '
        'corporation with its principal offices at 500 Innovation Drive, Suite 400, Austin, Texas '
        '78701 (the "Service Provider"), and Brightfield Holdings, LLC, a New York limited liability '
        'company with its principal offices at 200 Park Avenue, 30th Floor, New York, New York 10166 '
        '(the "Client").'
    )
    run.font.size = Pt(11)

    recitals_heading = doc.add_heading('RECITALS', level=1)
    recitals = [
        'WHEREAS, the Service Provider is engaged in the business of providing technology consulting, '
        'software development, and managed IT services to enterprise clients;',
        'WHEREAS, the Client desires to engage the Service Provider to provide certain technology '
        'services and Deliverables as more particularly described in individual Statements of Work;',
        'WHEREAS, the Parties desire to set forth the general terms and conditions that shall govern '
        'the provision of such services;',
        'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, '
        'and for other good and valuable consideration, the receipt and sufficiency of which are '
        'hereby acknowledged, the Parties agree as follows:',
    ]
    for text in recitals:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article I: Definitions ---
    doc.add_heading('ARTICLE I: DEFINITIONS', level=1)
    intro = doc.add_paragraph(
        'As used in this Agreement, the following terms shall have the meanings set forth below:'
    )
    intro.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for r in intro.runs:
        r.font.size = Pt(11)

    for i, (term, definition) in enumerate(DEFINED_TERMS, 1):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)

        num_run = p.add_run(f'1.{i}  ')
        num_run.font.size = Pt(11)

        term_run = p.add_run(f'"{term}"')
        term_run.bold = True
        term_run.font.size = Pt(11)

        def_run = p.add_run(f' means {definition}')
        def_run.font.size = Pt(11)

    # --- Article II: Scope of Services ---
    doc.add_heading('ARTICLE II: SCOPE OF SERVICES', level=1)
    scope_paras = [
        '2.1  Engagement. The Client hereby engages the Service Provider, and the Service Provider '
        'hereby accepts such engagement, to provide technology consulting, software development, '
        'and related services as described in individual Statements of Work executed by the Parties.',
        '2.2  Statements of Work. Each Statement of Work shall specify the scope of services, '
        'Deliverables, timeline, Key Personnel, acceptance criteria, and fees applicable to the '
        'particular engagement. In the event of a conflict between this Agreement and a Statement '
        'of Work, the terms of this Agreement shall control unless the Statement of Work expressly '
        'states otherwise.',
        '2.3  Change Orders. Either Party may request changes to a Statement of Work by submitting '
        'a written change order. No change order shall be effective unless signed by authorized '
        'representatives of both Parties. The Service Provider shall provide a revised Fee Schedule '
        'reflecting any additional costs associated with the change order.',
    ]
    for text in scope_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article III: Compensation ---
    doc.add_heading('ARTICLE III: COMPENSATION AND PAYMENT', level=1)
    comp_paras = [
        '3.1  Fees. The Client shall pay the Service Provider the fees set forth in the applicable '
        'Fee Schedule and Statement of Work. All fees are stated in United States Dollars and are '
        'exclusive of applicable taxes.',
        '3.2  Invoicing. The Service Provider shall submit invoices on a monthly basis for services '
        'rendered during the preceding calendar month. Each invoice shall include a detailed '
        'description of the services performed, hours expended by each member of Key Personnel, '
        'and any reimbursable expenses.',
        '3.3  Payment Terms. The Client shall pay all undisputed amounts within thirty (30) days '
        'of receipt of a proper invoice. Late payments shall accrue interest at a rate of one and '
        'one-half percent (1.5%) per month, or the maximum rate permitted by Applicable Law, '
        'whichever is less.',
        '3.4  Expense Reimbursement. The Service Provider shall be entitled to reimbursement for '
        'reasonable, pre-approved travel and out-of-pocket expenses incurred in connection with the '
        'performance of services. All expense claims must be accompanied by appropriate documentation.',
    ]
    for text in comp_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article IV: Intellectual Property ---
    doc.add_heading('ARTICLE IV: INTELLECTUAL PROPERTY RIGHTS', level=1)
    ip_paras = [
        '4.1  Work Product Ownership. All Deliverables and work product created by the Service '
        'Provider specifically for the Client under a Statement of Work shall be considered '
        'works made for hire and shall be the exclusive property of the Client.',
        '4.2  Pre-Existing IP. Each Party shall retain all right, title, and interest in and to '
        'its pre-existing Intellectual Property. The Service Provider hereby grants the Client a '
        'non-exclusive, perpetual, royalty-free license to use any pre-existing Intellectual Property '
        'incorporated into the Deliverables.',
        '4.3  Third-Party Materials. The Service Provider shall not incorporate any third-party '
        'materials into the Deliverables without the prior written consent of the Client, and shall '
        'ensure that appropriate licenses are obtained for any such materials.',
    ]
    for text in ip_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article V: Subcontracting ---
    doc.add_heading('ARTICLE V: SUBCONTRACTING', level=1)
    sub_paras = [
        '5.1  Prior Approval. The Service Provider shall not subcontract any portion of the services '
        'without the prior written approval of the Client. Any approved subcontractor shall be deemed '
        'a Permitted Subcontractor.',
        '5.2  Responsibility. The Service Provider shall remain fully responsible for the performance '
        'of all Permitted Subcontractors and shall ensure that each Permitted Subcontractor is bound '
        'by confidentiality and Intellectual Property obligations no less protective than those '
        'contained in this Agreement.',
    ]
    for text in sub_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article VI: Confidentiality ---
    doc.add_heading('ARTICLE VI: CONFIDENTIALITY', level=1)
    conf_paras = [
        '6.1  Obligations. Each Party agrees to hold the other Party\'s Confidential Information '
        'in strict confidence, to use such Confidential Information solely for the purposes of '
        'this Agreement, and to disclose such Confidential Information only to those employees '
        'and Permitted Subcontractors who have a need to know.',
        '6.2  Exclusions. Confidential Information shall not include information that: (a) is or '
        'becomes publicly available through no fault of the receiving Party; (b) was rightfully '
        'in the possession of the receiving Party prior to disclosure; (c) is independently developed '
        'by the receiving Party without use of the disclosing Party\'s Confidential Information; or '
        '(d) is rightfully obtained from a third party without restriction on disclosure.',
        '6.3  Duration. The obligations of confidentiality shall survive the termination or '
        'expiration of this Agreement for a period of five (5) years.',
    ]
    for text in conf_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article VII: Indemnification ---
    doc.add_heading('ARTICLE VII: INDEMNIFICATION', level=1)
    indem_paras = [
        '7.1  Indemnification by Service Provider. The Service Provider shall indemnify, defend, '
        'and hold harmless the Client and its officers, directors, employees, and agents (each an '
        'Indemnified Party) from and against any Claim and all associated Damages arising out of '
        'or relating to: (a) any breach of the Service Provider\'s representations, warranties, or '
        'obligations under this Agreement; (b) the negligence or willful misconduct of the Service '
        'Provider or its personnel; or (c) any infringement of third-party Intellectual Property rights.',
        '7.2  Indemnification by Client. The Client shall indemnify, defend, and hold harmless the '
        'Service Provider from and against any Claim and all associated Damages arising out of '
        'the Client\'s breach of its obligations under this Agreement or the Client\'s negligence '
        'or willful misconduct.',
        '7.3  Limitation of Liability. In no event shall either Party\'s aggregate liability exceed '
        'the Liability Cap. Neither Party shall be liable for any indirect, incidental, consequential, '
        'special, or punitive Damages, regardless of the form of action.',
    ]
    for text in indem_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article VIII: Representations and Warranties ---
    doc.add_heading('ARTICLE VIII: REPRESENTATIONS AND WARRANTIES', level=1)
    rw_paras = [
        '8.1  Mutual Representations. Each Party represents and warrants that: (a) it is duly '
        'organized and validly existing under the laws of its jurisdiction of organization; (b) it '
        'has full power and authority to enter into this Agreement; and (c) the execution and '
        'performance of this Agreement does not violate any Applicable Law or any other agreement '
        'to which it is a party.',
        '8.2  Service Provider Warranties. The Service Provider represents and warrants that: '
        '(a) all services shall be performed in a professional and workmanlike manner by qualified '
        'Key Personnel; (b) all Deliverables shall conform to the specifications set forth in the '
        'applicable Statement of Work; and (c) the Service Provider shall comply with the Service '
        'Level Agreement.',
    ]
    for text in rw_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article IX: Termination ---
    doc.add_heading('ARTICLE IX: TERM AND TERMINATION', level=1)
    term_paras = [
        '9.1  Term. This Agreement shall remain in effect for the Term, and shall automatically '
        'renew for successive one (1) year periods unless either Party provides written Notice of '
        'non-renewal at least sixty (60) days prior to the end of the then-current term.',
        '9.2  Termination for Material Breach. Either Party may terminate this Agreement upon '
        'written Notice if the other Party commits a Material Breach that remains uncured after '
        'the applicable cure period.',
        '9.3  Termination for Convenience. Either Party may exercise its right of Termination '
        'for Convenience in accordance with Section 1.25 of this Agreement.',
        '9.4  Effect of Termination. Upon termination or expiration of this Agreement, the Service '
        'Provider shall deliver all completed and in-progress Deliverables to the Client, and the '
        'Client shall pay for all services rendered through the date of termination.',
    ]
    for text in term_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article X: Force Majeure ---
    doc.add_heading('ARTICLE X: FORCE MAJEURE', level=1)
    fm_paras = [
        '10.1  Excuse of Performance. Neither Party shall be liable for any failure or delay in '
        'the performance of its obligations under this Agreement to the extent such failure or delay '
        'is caused by a Force Majeure Event, provided that the affected Party gives prompt Notice '
        'to the other Party and uses commercially reasonable efforts to mitigate the effects thereof.',
        '10.2  Extended Force Majeure. If a Force Majeure Event continues for more than ninety (90) '
        'consecutive Business Days, either Party may terminate this Agreement upon thirty (30) days\' '
        'written Notice to the other Party.',
    ]
    for text in fm_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article XI: Dispute Resolution ---
    doc.add_heading('ARTICLE XI: DISPUTE RESOLUTION', level=1)
    dr_paras = [
        '11.1  Negotiation. The Parties shall attempt in good faith to resolve any dispute arising '
        'out of or relating to this Agreement through negotiation between their respective senior '
        'executives within thirty (30) days after one Party delivers Notice of the dispute.',
        '11.2  Mediation. If the dispute is not resolved through negotiation, the Parties agree to '
        'submit the dispute to non-binding mediation administered by the American Arbitration '
        'Association before initiating any litigation.',
        '11.3  Governing Law. This Agreement shall be governed by and construed in accordance with '
        'the Governing Law. Any litigation arising under this Agreement shall be brought exclusively '
        'in the state or federal courts located in Wilmington, Delaware.',
    ]
    for text in dr_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Article XII: General Provisions ---
    doc.add_heading('ARTICLE XII: GENERAL PROVISIONS', level=1)
    gp_paras = [
        '12.1  Notices. All Notices required or permitted under this Agreement shall be in writing '
        'and shall be deemed given when delivered personally, sent by overnight courier (with '
        'confirmation of delivery), sent by certified mail (return receipt requested), or sent by '
        'electronic mail with confirmation of receipt, to the addresses set forth on the signature '
        'page hereof.',
        '12.2  Assignment. Neither Party may assign this Agreement or any rights or obligations '
        'hereunder without the prior written consent of the other Party, except that either Party '
        'may assign this Agreement to an Affiliate or in connection with a Change of Control.',
        '12.3  Entire Agreement. This Agreement, together with all Exhibits, Schedules, and '
        'Statements of Work, constitutes the entire agreement between the Parties with respect to '
        'the subject matter hereof and supersedes all prior agreements and understandings.',
        '12.4  Amendments. No amendment or modification of this Agreement shall be effective unless '
        'made in writing and signed by authorized representatives of both Parties.',
        '12.5  Severability. If any provision of this Agreement is held to be invalid or '
        'unenforceable by a court of competent jurisdiction, the remaining provisions shall continue '
        'in full force and effect.',
        '12.6  Waiver. The failure of either Party to enforce any right or provision of this '
        'Agreement shall not constitute a waiver of such right or provision.',
    ]
    for text in gp_paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # --- Signature Block ---
    doc.add_paragraph()  # spacer
    sig_header = doc.add_paragraph('IN WITNESS WHEREOF, the Parties have executed this Agreement '
                                   'as of the Effective Date.')
    sig_header.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for r in sig_header.runs:
        r.font.size = Pt(11)

    doc.add_paragraph()
    for entity, name, title_text in [
        ('MERIDIAN TECHNOLOGY SOLUTIONS, INC.', 'Jonathan R. Mitchell', 'Chief Executive Officer'),
        ('BRIGHTFIELD HOLDINGS, LLC', 'Catherine A. Donovan', 'Managing Director'),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(entity)
        run.bold = True
        run.font.size = Pt(11)

        doc.add_paragraph()
        p2 = doc.add_paragraph()
        p2.add_run('By: ________________________________').font.size = Pt(11)
        p3 = doc.add_paragraph()
        p3.add_run(f'Name: {name}').font.size = Pt(11)
        p4 = doc.add_paragraph()
        p4.add_run(f'Title: {title_text}').font.size = Pt(11)
        p5 = doc.add_paragraph()
        p5.add_run('Date: ________________________________').font.size = Pt(11)
        doc.add_paragraph()

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


create_initial()
