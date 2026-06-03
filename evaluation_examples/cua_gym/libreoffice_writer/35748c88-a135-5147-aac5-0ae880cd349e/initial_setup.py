"""
Initial Setup: Business document with default toolbar configuration
Task ID: writer_biz_079
Domain: libreoffice_writer

Creates a realistic business memo document and opens it in LibreOffice Writer
with the default toolbar configuration (no custom toolbar).
"""

import os
import shlex
import subprocess
import time

WORKDIR = '/home/user'
TASK_ID = 'writer_biz_079'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Company Header ---
    heading = doc.add_heading('Meridian Consulting Group', level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # --- Memo Line ---
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle.add_run('INTERNAL MEMORANDUM')
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    # --- Horizontal Rule ---
    doc.add_paragraph('_' * 72)

    # --- Memo Metadata ---
    meta_fields = [
        ('TO:', 'Senior Leadership Team'),
        ('FROM:', 'Rachel Torres, Director of Operations'),
        ('DATE:', 'March 28, 2026'),
        ('RE:', 'Q1 2026 Operational Performance Review'),
    ]
    for label, value in meta_fields:
        p = doc.add_paragraph()
        run_label = p.add_run(label + '  ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_value = p.add_run(value)
        run_value.font.size = Pt(11)

    doc.add_paragraph('_' * 72)

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'This memorandum provides a comprehensive review of Meridian Consulting '
        'Group\'s operational performance during Q1 2026. Overall, the firm has '
        'demonstrated strong growth in client acquisition and project delivery, '
        'while maintaining cost discipline across all departments.'
    )

    # --- Financial Highlights ---
    doc.add_heading('Financial Highlights', level=2)
    doc.add_paragraph(
        'Revenue for Q1 2026 reached $4.87 million, representing a 12.3% increase '
        'over Q1 2025. Operating margins improved to 23.1%, up from 19.8% in the '
        'prior year period. Key drivers include:'
    )
    bullet_items = [
        'Strategic consulting engagements: $2.14M (+18.5%)',
        'Technology implementation services: $1.62M (+9.2%)',
        'Managed services contracts: $0.73M (+4.1%)',
        'Training and development programs: $0.38M (+22.7%)',
    ]
    for item in bullet_items:
        doc.add_paragraph(item, style='List Bullet')

    # --- Client Portfolio ---
    doc.add_heading('Client Portfolio Update', level=2)
    doc.add_paragraph(
        'During Q1, we onboarded 14 new clients across financial services, '
        'healthcare, and technology sectors. Notable new engagements include:'
    )

    # Table of new clients
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['Client', 'Industry', 'Engagement Type', 'Contract Value']
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    clients = [
        ['Apex Financial Partners', 'Financial Services', 'Digital Transformation', '$420,000'],
        ['NovaCare Health Systems', 'Healthcare', 'Process Optimization', '$285,000'],
        ['TerraForge Technologies', 'Technology', 'Cloud Migration', '$510,000'],
        ['Pinnacle Insurance Group', 'Insurance', 'Regulatory Compliance', '$195,000'],
        ['Coastal Manufacturing Inc.', 'Manufacturing', 'Supply Chain Analysis', '$340,000'],
    ]
    for r, row_data in enumerate(clients, 1):
        for c, val in enumerate(row_data):
            table.cell(r, c).text = val

    # --- Team Performance ---
    doc.add_heading('Team Performance Metrics', level=2)
    doc.add_paragraph(
        'Employee utilization rates averaged 78.4% across the consulting practice, '
        'exceeding our target of 75%. The technology services team achieved 82.1% '
        'utilization, the highest in firm history for a Q1 period.'
    )
    doc.add_paragraph(
        'Staff retention remained strong at 94.2%, with voluntary turnover limited '
        'to 3 departures out of 52 full-time consultants. We completed 8 new hires '
        'during the quarter, bringing total headcount to 57.'
    )

    # --- Operational Challenges ---
    doc.add_heading('Operational Challenges', level=2)
    numbered_items = [
        'Project delivery timelines extended by an average of 1.3 weeks due to '
        'client-side resource constraints at two major engagements.',
        'Technology infrastructure costs increased 8.4% due to expanded cloud '
        'hosting requirements for the managed services division.',
        'Recruitment of senior-level consultants with specialized healthcare IT '
        'expertise continues to be challenging in the current labor market.',
    ]
    for item in numbered_items:
        doc.add_paragraph(item, style='List Number')

    # --- Looking Ahead ---
    doc.add_heading('Q2 2026 Outlook', level=2)
    doc.add_paragraph(
        'The sales pipeline for Q2 is robust, with $3.2 million in qualified '
        'opportunities at various stages of negotiation. Key priorities for the '
        'upcoming quarter include:'
    )
    q2_items = [
        'Finalize partnership agreement with DataBridge Analytics for joint '
        'service offerings in the financial services vertical.',
        'Launch the revised consultant development program with emphasis on '
        'cloud architecture and data analytics certifications.',
        'Complete the office expansion to accommodate projected headcount growth '
        'through year-end.',
        'Implement the new project management platform (Meridian ProjectHub) '
        'across all practice areas by end of April.',
    ]
    for item in q2_items:
        doc.add_paragraph(item, style='List Bullet')

    # --- Signature ---
    doc.add_paragraph('')
    sig = doc.add_paragraph()
    sig_run = sig.add_run('Rachel Torres')
    sig_run.bold = True
    sig_run.font.size = Pt(11)
    title_p = doc.add_paragraph('Director of Operations')
    title_p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph('Meridian Consulting Group')
    doc.add_paragraph('rachel.torres@meridiancg.com')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # --- Ensure no custom toolbars exist (clean default state) ---
    toolbar_dir = '/home/user/.config/libreoffice/4/user/config/soffice.cfg/modules/swriter/toolbar'
    # Remove any custom toolbar files that might exist
    import glob
    for f in glob.glob(os.path.join(toolbar_dir, 'custom_toolbar_*.xml')):
        os.remove(f)
        print(f'Removed stale custom toolbar: {f}')

    # Launch Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
