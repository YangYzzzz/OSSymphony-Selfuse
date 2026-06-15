"""
Reward Script: HR Business Continuity Plan
Task ID: writer_hr_082
Domain: libreoffice_writer
Scoring:
  Component 1: Heading hierarchy (4 levels used) - 0.20 pts
  Component 2: Tables present (multiple tables for data) - 0.20 pts
  Component 3: Critical function priority table (10 ranked functions) - 0.20 pts
  Component 4: TOC field code present - 0.15 pts
  Component 5: Contact tree structure present - 0.10 pts
  Component 6: Revision log table present - 0.15 pts
"""

import os
from docx import Document
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_082'


def persist_app_state(domain: str):
    """Best-effort save of any open LibreOffice document."""
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            import time
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Component 1: Heading hierarchy - all 4 levels used (0.20 points)
    # Initial file has NO headings (all Normal style). Golden has Heading 1-4.
    try:
        heading_levels_found = set()
        for p in doc.paragraphs:
            sname = p.style.name if p.style else ''
            if sname == 'Heading 1':
                heading_levels_found.add(1)
            elif sname == 'Heading 2':
                heading_levels_found.add(2)
            elif sname == 'Heading 3':
                heading_levels_found.add(3)
            elif sname == 'Heading 4':
                heading_levels_found.add(4)

        # Need all 4 levels for full points
        levels_present = len(heading_levels_found.intersection({1, 2, 3, 4}))
        if levels_present == 4:
            print(f"PASS: Component 1 - All 4 heading levels found: {sorted(heading_levels_found)} (0.20 pts)")
            total_score += 0.20
        elif levels_present >= 3:
            partial = 0.15
            print(f"PARTIAL: Component 1 - {levels_present}/4 heading levels found: {sorted(heading_levels_found)} ({partial} pts)")
            total_score += partial
        elif levels_present >= 2:
            partial = 0.10
            print(f"PARTIAL: Component 1 - {levels_present}/4 heading levels found: {sorted(heading_levels_found)} ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 - Only {levels_present}/4 heading levels found: {sorted(heading_levels_found)}")
    except Exception as e:
        print(f"ERROR: Component 1 - {e}")

    # Component 2: Multiple tables present for structured data (0.20 points)
    # Initial file has 0 tables. Golden has 13 tables.
    try:
        num_tables = len(doc.tables)
        if num_tables >= 8:
            print(f"PASS: Component 2 - {num_tables} tables found (need >= 8 for full credit) (0.20 pts)")
            total_score += 0.20
        elif num_tables >= 5:
            partial = 0.12
            print(f"PARTIAL: Component 2 - {num_tables} tables found (need >= 8 for full) ({partial} pts)")
            total_score += partial
        elif num_tables >= 2:
            partial = 0.06
            print(f"PARTIAL: Component 2 - {num_tables} tables found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 2 - Only {num_tables} tables found (need >= 8)")
    except Exception as e:
        print(f"ERROR: Component 2 - {e}")

    # Component 3: Critical function priority table with 10 ranked functions (0.20 points)
    # Initial file has no tables. Golden has a 6-col table with Priority/Function/RTO/RPO/Impact Level/Dependencies
    try:
        priority_table_found = False
        priority_rows = 0
        for t in doc.tables:
            if len(t.rows) >= 2 and len(t.columns) >= 4:
                header_texts = [c.text.strip().lower() for c in t.rows[0].cells]
                # Look for priority/function/RTO columns
                has_priority = any('priority' in h or 'rank' in h for h in header_texts)
                has_function = any('function' in h for h in header_texts)
                has_rto = any('rto' in h or 'recovery' in h for h in header_texts)
                if has_priority and has_function and has_rto:
                    priority_table_found = True
                    priority_rows = len(t.rows) - 1  # exclude header
                    break

        if priority_table_found and priority_rows >= 10:
            print(f"PASS: Component 3 - Critical function priority table found with {priority_rows} functions (0.20 pts)")
            total_score += 0.20
        elif priority_table_found and priority_rows >= 5:
            partial = 0.12
            print(f"PARTIAL: Component 3 - Priority table found but only {priority_rows}/10 functions ({partial} pts)")
            total_score += partial
        elif priority_table_found:
            partial = 0.06
            print(f"PARTIAL: Component 3 - Priority table found but only {priority_rows} functions ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 3 - No critical function priority table found")
    except Exception as e:
        print(f"ERROR: Component 3 - {e}")

    # Component 4: Table of Contents with field code (0.15 points)
    # Initial file has no TOC. Golden has TOC field in XML.
    try:
        body_xml = doc.element.body.xml
        has_toc_field = 'TOC' in body_xml and 'instrText' in body_xml
        # Also check if there is a Heading 1 paragraph mentioning TOC
        has_toc_heading = False
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith('Heading'):
                if 'table of contents' in p.text.lower() or 'toc' in p.text.lower():
                    has_toc_heading = True
                    break

        if has_toc_field:
            print(f"PASS: Component 4 - TOC field code found in document XML (0.15 pts)")
            total_score += 0.15
        elif has_toc_heading:
            partial = 0.08
            print(f"PARTIAL: Component 4 - TOC heading found but no field code ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 4 - No TOC field or TOC heading found")
    except Exception as e:
        print(f"ERROR: Component 4 - {e}")

    # Component 5: Emergency contact tree structure (0.10 points)
    # Initial file has no structured contact tree. Golden has text-based tree with unicode chars
    # and tables for manager/team contacts.
    try:
        tree_chars_found = False
        tree_contact_found = False
        for p in doc.paragraphs:
            text = p.text
            # Unicode box-drawing characters indicate tree diagram
            if any(ch in text for ch in ['│', '┌', '┼', '├', '└', '┐', '─']):
                tree_chars_found = True
            # Check for VP contact info (Natasha Volkov)
            if 'volkov' in text.lower() and ('ext' in text.lower() or 'email' in text.lower() or '@' in text):
                tree_contact_found = True

        # Also check for contact-related tables (manager/team lead tables)
        contact_tables = 0
        for t in doc.tables:
            if len(t.rows) >= 2:
                header_texts = [c.text.strip().lower() for c in t.rows[0].cells]
                if any('manager' in h or 'team lead' in h or 'reports to' in h or 'extension' in h for h in header_texts):
                    contact_tables += 1

        # Require STRUCTURAL elements (tree chars or contact tables), not just name mentions
        has_structure = tree_chars_found or contact_tables >= 1
        if has_structure and tree_contact_found:
            print(f"PASS: Component 5 - Contact tree structure found (tree_chars={tree_chars_found}, tables={contact_tables}, contact={tree_contact_found}) (0.10 pts)")
            total_score += 0.10
        elif has_structure:
            partial = 0.05
            print(f"PARTIAL: Component 5 - Structural tree elements found but missing VP contact (tree_chars={tree_chars_found}, tables={contact_tables}) ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 5 - No contact tree structure found")
    except Exception as e:
        print(f"ERROR: Component 5 - {e}")

    # Component 6: Revision log table (0.15 points)
    # Initial file has no tables. Golden has a version/date/author/description revision log table.
    try:
        revision_table_found = False
        revision_rows = 0
        for t in doc.tables:
            if len(t.rows) >= 2:
                header_texts = [c.text.strip().lower() for c in t.rows[0].cells]
                has_version = any('version' in h for h in header_texts)
                has_date = any('date' in h for h in header_texts)
                has_author = any('author' in h for h in header_texts)
                if has_version and has_date and has_author:
                    revision_table_found = True
                    revision_rows = len(t.rows) - 1  # exclude header
                    break

        if revision_table_found and revision_rows >= 3:
            print(f"PASS: Component 6 - Revision log table found with {revision_rows} entries (0.15 pts)")
            total_score += 0.15
        elif revision_table_found:
            partial = 0.08
            print(f"PARTIAL: Component 6 - Revision log found but only {revision_rows} entries ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 6 - No revision log table found")
    except Exception as e:
        print(f"ERROR: Component 6 - {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
