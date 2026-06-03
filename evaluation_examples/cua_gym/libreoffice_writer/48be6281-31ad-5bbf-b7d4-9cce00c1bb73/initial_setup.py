"""
Initial Setup: HR Business Continuity Plan - rough notes document
Task ID: writer_hr_082
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_082'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Title
    title = doc.add_heading('HR Business Continuity Plan', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle / meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run('Meridian Technologies Inc.')
    run.font.size = Pt(14)
    run.bold = True
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = meta2.add_run('Human Resources Department')
    run2.font.size = Pt(12)

    doc.add_paragraph('')  # spacer

    # Rough notes section
    doc.add_paragraph('--- ROUGH NOTES AND PLANNING ---').bold = True

    doc.add_paragraph(
        'This document contains preliminary notes for the HR Business Continuity Plan. '
        'The following sections need to be developed into a comprehensive plan with '
        'proper formatting, tables, and organizational hierarchy.'
    )

    doc.add_paragraph('')

    # Section: Document Control notes
    p = doc.add_paragraph('DOCUMENT CONTROL')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Need to create a document control page with version history. '
        'Current version is 1.0, drafted by Patricia Sandoval, HR Director. '
        'Approved by Raymond Kowalski, Chief Operating Officer. '
        'Effective date: January 15, 2026. Next review scheduled for July 2026.'
    )

    doc.add_paragraph('')

    # Section: Emergency Contacts notes
    p = doc.add_paragraph('EMERGENCY CONTACT TREE')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Need to build an org-chart style contact tree. VP of HR Natasha Volkov (ext 2001) '
        'leads. Under her: Director of Talent Acquisition Derek Huang (ext 2010), '
        'Director of Employee Relations Samira Okafor (ext 2020), '
        'Director of Compensation & Benefits Luis Restrepo (ext 2030). '
        'Each director oversees 2-3 managers who in turn lead teams. '
        'Should use text boxes or shapes with connecting lines to show hierarchy.'
    )

    doc.add_paragraph('')

    # Section: Critical Functions notes
    p = doc.add_paragraph('CRITICAL FUNCTION PRIORITIES')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Identify and rank 10 critical HR functions by recovery time objective (RTO). '
        'Functions include: Payroll Processing (RTO 4hrs), Benefits Administration (RTO 8hrs), '
        'Employee Safety & Compliance (RTO 4hrs), Recruitment Pipeline (RTO 24hrs), '
        'HRIS System Access (RTO 2hrs), Workers Comp Claims (RTO 8hrs), '
        'Training & Development (RTO 72hrs), Performance Management (RTO 48hrs), '
        'Employee Relations (RTO 12hrs), Onboarding (RTO 24hrs). '
        'Create a table ranked by priority with impact assessment.'
    )

    doc.add_paragraph('')

    # Section: Alternate Work notes
    p = doc.add_paragraph('ALTERNATE WORK ARRANGEMENTS')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Document remote work activation procedures, alternate site locations, '
        'equipment requirements, VPN access protocols, and team rotation schedules. '
        'Primary alternate site: Meridian Westside Campus, 450 Industrial Blvd. '
        'Secondary: Co-working agreement with FlexSpace Downtown. '
        'Include hot-desking policy and equipment checkout procedures.'
    )

    doc.add_paragraph('')

    # Section: Communication Protocols notes
    p = doc.add_paragraph('COMMUNICATION PROTOCOLS')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Need two tables: internal communication (HR staff, department heads, all employees) '
        'and external communication (vendors, insurance carriers, government agencies, '
        'temp staffing agencies). Define channels, responsible parties, timing, and escalation paths.'
    )

    doc.add_paragraph('')

    # Section: Resource Requirements notes
    p = doc.add_paragraph('RESOURCE REQUIREMENTS')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Catalog all resources needed for BCP activation: technology (laptops, VPN tokens, '
        'backup servers), personnel (cross-trained staff, temp agency contacts), '
        'financial (emergency petty cash, vendor payment authorization), '
        'physical (office supplies, printed forms, first aid kits). '
        'Include cost estimates and procurement timelines.'
    )

    doc.add_paragraph('')

    # Section: Testing Schedule notes
    p = doc.add_paragraph('TESTING SCHEDULE')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Plan quarterly testing activities: tabletop exercises, communication drills, '
        'system failover tests, and full-scale simulations. '
        'Track test dates, participants, findings, and corrective actions.'
    )

    doc.add_paragraph('')

    # Section: Revision Log notes
    p = doc.add_paragraph('REVISION LOG')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Track all changes: Version 0.1 initial draft by Patricia Sandoval (Nov 2025), '
        'Version 0.5 review by legal team (Dec 2025), '
        'Version 0.9 executive feedback incorporated (Jan 2026), '
        'Version 1.0 final approved (Jan 15, 2026).'
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'TODO: Format all sections with proper heading hierarchy (Heading 1 through Heading 4), '
        'create tables for data, build contact tree diagram, and generate Table of Contents.'
    )

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
