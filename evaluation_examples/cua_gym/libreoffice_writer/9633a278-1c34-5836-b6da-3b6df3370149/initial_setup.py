"""
Initial Setup: Remove all hyperlinks from a contract document
Task ID: writer_legal_015
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_015'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph with blue underlined text."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {
        qn('r:id'): r_id,
    })
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})

    # Blue color
    color_elem = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0563C1'})
    rPr.append(color_elem)

    # Underline
    u_elem = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr.append(u_elem)

    # Font size 11pt
    sz_elem = paragraph._element.makeelement(qn('w:sz'), {qn('w:val'): '22'})
    rPr.append(sz_elem)

    new_run.append(rPr)
    t_elem = paragraph._element.makeelement(qn('w:t'), {})
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    new_run.append(t_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_normal_run(paragraph, text, bold=False, size_pt=11):
    """Add a normal run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    return run


def create_initial():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    title = doc.add_heading('SOFTWARE LICENSING AGREEMENT', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Parties ---
    p = doc.add_paragraph()
    add_normal_run(p, 'This Software Licensing Agreement ("Agreement") is entered into as of March 15, 2025, by and between ')
    add_normal_run(p, 'Meridian Technologies, Inc.', bold=True)
    add_normal_run(p, ', a Delaware corporation with its principal place of business at 2400 Innovation Drive, Suite 800, San Jose, CA 95134 ("Licensor"), and ')
    add_normal_run(p, 'Oakridge Financial Services, LLC', bold=True)
    add_normal_run(p, ', a New York limited liability company with offices at 350 Park Avenue, 22nd Floor, New York, NY 10022 ("Licensee").')

    doc.add_paragraph()

    # --- Section 1: Definitions ---
    h1 = doc.add_heading('1. DEFINITIONS', level=1)

    p = doc.add_paragraph()
    add_normal_run(p, '1.1 "Software" means the Meridian Analytics Platform, version 4.2, including all modules, updates, patches, and documentation provided by Licensor during the term of this Agreement.')

    p = doc.add_paragraph()
    add_normal_run(p, '1.2 "Authorized Users" means employees, contractors, and agents of Licensee who have been granted access to the Software in accordance with Section 3 of this Agreement and applicable data protection regulations as defined under ')
    # Hyperlink 1
    add_hyperlink(p, 'https://www.congress.gov/bill/116th-congress/senate-bill/2968', 'the California Consumer Privacy Act (CCPA)')
    add_normal_run(p, '.')

    p = doc.add_paragraph()
    add_normal_run(p, '1.3 "Confidential Information" means all proprietary data, trade secrets, and business information disclosed by either party, subject to the protections outlined in ')
    # Hyperlink 2
    add_hyperlink(p, 'https://www.law.cornell.edu/uscode/text/18/1836', '18 U.S.C. \u00a7 1836 (Defend Trade Secrets Act)')
    add_normal_run(p, '.')

    doc.add_paragraph()

    # --- Section 2: Grant of License ---
    h2 = doc.add_heading('2. GRANT OF LICENSE', level=1)

    p = doc.add_paragraph()
    add_normal_run(p, '2.1 Subject to the terms and conditions of this Agreement and the payment of all applicable fees, Licensor hereby grants to Licensee a non-exclusive, non-transferable, limited license to install and use the Software on up to fifty (50) workstations within Licensee\'s facilities.')

    p = doc.add_paragraph()
    add_normal_run(p, '2.2 Licensee shall comply with all applicable export control laws and regulations, including those set forth in ')
    # Hyperlink 3
    add_hyperlink(p, 'https://www.ecfr.gov/current/title-15/subtitle-B/chapter-VII/subchapter-C/part-730', 'the Export Administration Regulations (15 CFR Part 730-774)')
    add_normal_run(p, ', and shall not export or re-export the Software to any prohibited jurisdiction without prior written authorization from the appropriate governmental authority.')

    doc.add_paragraph()

    # --- Section 3: Intellectual Property ---
    h3 = doc.add_heading('3. INTELLECTUAL PROPERTY', level=1)

    p = doc.add_paragraph()
    add_normal_run(p, '3.1 All intellectual property rights in and to the Software, including all patents, copyrights, trademarks, and trade secrets, shall remain the exclusive property of Licensor. This Agreement does not convey any ownership interest in the Software to Licensee.')

    p = doc.add_paragraph()
    add_normal_run(p, '3.2 The Software is protected under ')
    # Hyperlink 4
    add_hyperlink(p, 'https://www.law.cornell.edu/uscode/text/17/chapter-1', 'Title 17 of the United States Code (Copyright Act)')
    add_normal_run(p, ' and applicable international treaties including the ')
    # Hyperlink 5
    add_hyperlink(p, 'https://www.wipo.int/treaties/en/ip/berne/', 'Berne Convention for the Protection of Literary and Artistic Works')
    add_normal_run(p, '. Unauthorized reproduction or distribution may result in civil and criminal penalties.')

    doc.add_paragraph()

    # --- Section 4: Data Protection ---
    h4 = doc.add_heading('4. DATA PROTECTION AND PRIVACY', level=1)

    p = doc.add_paragraph()
    add_normal_run(p, '4.1 Licensor shall process any personal data provided by Licensee in accordance with ')
    # Hyperlink 6
    add_hyperlink(p, 'https://gdpr-info.eu/', 'the General Data Protection Regulation (GDPR)')
    add_normal_run(p, ' where applicable, and shall implement appropriate technical and organizational measures to ensure the security of such data.')

    p = doc.add_paragraph()
    add_normal_run(p, '4.2 In the event of a data breach affecting Licensee\'s personal data, Licensor shall notify Licensee within seventy-two (72) hours of becoming aware of the breach, in compliance with applicable notification requirements under state and federal law.')

    doc.add_paragraph()

    # --- Section 5: Limitation of Liability ---
    h5 = doc.add_heading('5. LIMITATION OF LIABILITY', level=1)

    p = doc.add_paragraph()
    add_normal_run(p, '5.1 TO THE MAXIMUM EXTENT PERMITTED BY APPLICABLE LAW, INCLUDING ')
    # Hyperlink 7
    add_hyperlink(p, 'https://www.law.cornell.edu/ucc/article2', 'the Uniform Commercial Code (Article 2)')
    add_normal_run(p, ', IN NO EVENT SHALL LICENSOR BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF OR RELATED TO THIS AGREEMENT.')

    p = doc.add_paragraph()
    add_normal_run(p, '5.2 Licensor\'s total cumulative liability under this Agreement shall not exceed the total fees paid by Licensee during the twelve (12) month period immediately preceding the event giving rise to the claim.')

    doc.add_paragraph()

    # --- Section 6: Governing Law ---
    h6 = doc.add_heading('6. GOVERNING LAW AND DISPUTE RESOLUTION', level=1)

    p = doc.add_paragraph()
    add_normal_run(p, '6.1 This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware, without regard to its conflict of laws provisions, and in accordance with ')
    # Hyperlink 8
    add_hyperlink(p, 'https://www.law.cornell.edu/uscode/text/9', 'the Federal Arbitration Act (9 U.S.C. \u00a7\u00a7 1-16)')
    add_normal_run(p, '.')

    p = doc.add_paragraph()
    add_normal_run(p, '6.2 Any dispute arising under this Agreement that cannot be resolved through good-faith negotiation shall be submitted to binding arbitration administered by the American Arbitration Association in accordance with its Commercial Arbitration Rules. The arbitration shall take place in Wilmington, Delaware.')

    doc.add_paragraph()

    # --- Signature Block ---
    h7 = doc.add_heading('SIGNATURES', level=1)

    p = doc.add_paragraph()
    add_normal_run(p, 'IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')

    doc.add_paragraph()

    p = doc.add_paragraph()
    add_normal_run(p, 'LICENSOR: Meridian Technologies, Inc.', bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_normal_run(p, 'By: ________________________________')
    p = doc.add_paragraph()
    add_normal_run(p, 'Name: Jonathan R. Whitfield')
    p = doc.add_paragraph()
    add_normal_run(p, 'Title: Chief Executive Officer')
    p = doc.add_paragraph()
    add_normal_run(p, 'Date: March 15, 2025')

    doc.add_paragraph()

    p = doc.add_paragraph()
    add_normal_run(p, 'LICENSEE: Oakridge Financial Services, LLC', bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_normal_run(p, 'By: ________________________________')
    p = doc.add_paragraph()
    add_normal_run(p, 'Name: Patricia M. Sullivan')
    p = doc.add_paragraph()
    add_normal_run(p, 'Title: Managing Director')
    p = doc.add_paragraph()
    add_normal_run(p, 'Date: March 15, 2025')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
