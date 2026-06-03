"""
Reward Script: Remove all hyperlinks from contract document
Task ID: writer_legal_015
Domain: libreoffice_writer
Scoring:
  Component 1 (0.5): No hyperlink XML elements in document paragraphs
  Component 2 (0.2): No external hyperlink relationships in document rels
  Component 3 (0.3): Display text preserved AND no hyperlink-style formatting (blue color 0563C1)
"""

import os
from docx import Document
from docx.shared import RGBColor

WORKDIR = '/home/user'
TASK_ID = 'writer_legal_015'

# The 8 display texts that were originally hyperlinked
EXPECTED_TEXTS = [
    'the California Consumer Privacy Act (CCPA)',
    '18 U.S.C. \u00a7 1836 (Defend Trade Secrets Act)',
    'the Export Administration Regulations (15 CFR Part 730-774)',
    'Title 17 of the United States Code (Copyright Act)',
    'Berne Convention for the Protection of Literary and Artistic Works',
    'the General Data Protection Regulation (GDPR)',
    'the Uniform Commercial Code (Article 2)',
    'the Federal Arbitration Act (9 U.S.C. \u00a7\u00a7 1-16)',
]

HYPERLINK_BLUE = RGBColor(0x05, 0x63, 0xC1)


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Component 1: No hyperlink XML elements in paragraphs (0.5 points)
    # In initial_env there are 8 hyperlinks; in golden_env there should be 0.
    try:
        hyperlink_count = 0
        for para in doc.paragraphs:
            hyperlinks = para._element.findall('.//w:hyperlink', ns)
            hyperlink_count += len(hyperlinks)

        if hyperlink_count == 0:
            print(f"PASS: Component 1 — No hyperlink elements found (0.5 pts)")
            total_score += 0.5
        else:
            print(f"FAIL: Component 1 — Found {hyperlink_count} hyperlink elements (expected 0)")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: No external hyperlink relationships (0.2 points)
    # In initial_env there are 8 external hyperlink rels; in golden_env there should be 0.
    try:
        hyperlink_rels = 0
        for rid, rel in doc.part.rels.items():
            if 'hyperlink' in str(rel.reltype).lower():
                hyperlink_rels += 1

        if hyperlink_rels == 0:
            print(f"PASS: Component 2 — No external hyperlink relationships (0.2 pts)")
            total_score += 0.2
        else:
            print(f"FAIL: Component 2 — Found {hyperlink_rels} hyperlink relationships (expected 0)")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Display text preserved AND no hyperlink-style blue color (0.3 points)
    # In initial_env, the text exists but runs inside <w:hyperlink> have blue (0563C1) => FAIL.
    # In golden_env, the text exists as plain runs with no blue color => PASS.
    # NOTE: para.runs does NOT include runs inside <w:hyperlink> elements, so we must
    # scan ALL <w:r> elements in the body XML to detect blue-colored runs.
    try:
        from docx.text.run import Run
        full_text = ' '.join(p.text for p in doc.paragraphs)

        # Check all 8 display texts are preserved
        texts_found = 0
        for text in EXPECTED_TEXTS:
            if text in full_text:
                texts_found += 1
            else:
                print(f"  MISSING text: {text!r}")

        # Check no runs anywhere in the body have hyperlink-style blue color
        # This includes runs inside <w:hyperlink> elements that para.runs misses
        blue_runs_found = 0
        body = doc.element.body
        for r_elem in body.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            try:
                # Create a temporary Run to read font properties
                run = Run(r_elem, doc.paragraphs[0])
                if run.font.color and run.font.color.rgb == HYPERLINK_BLUE:
                    blue_runs_found += 1
            except Exception:
                pass

        if texts_found == len(EXPECTED_TEXTS) and blue_runs_found == 0:
            print(f"PASS: Component 3 — All {texts_found} display texts preserved, no hyperlink-blue runs (0.3 pts)")
            total_score += 0.3
        elif texts_found == len(EXPECTED_TEXTS) and blue_runs_found > 0:
            print(f"FAIL: Component 3 — Texts preserved but {blue_runs_found} runs still have hyperlink-blue color")
        else:
            print(f"FAIL: Component 3 — Only {texts_found}/{len(EXPECTED_TEXTS)} display texts found, {blue_runs_found} blue runs")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    final_score = min(total_score, 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Persistence hook: save any unsaved LibreOffice edits
def persist_app_state():
    import time
    os.environ["DISPLAY"] = ":0"
    try:
        import pyautogui
        pyautogui.hotkey("ctrl", "s")
        time.sleep(0.8)
        print("PERSIST: ctrl+s sent for libreoffice_writer")
    except Exception as e:
        print(f"PERSIST_WARN: save hook failed: {e}")


# Entry point
file_path = f'{WORKDIR}/{TASK_ID}.docx'
persist_app_state()

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
