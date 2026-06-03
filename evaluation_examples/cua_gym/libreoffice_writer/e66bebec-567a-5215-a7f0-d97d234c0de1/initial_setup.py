"""
Initial Setup: Written warning letter template with merge fields
Task ID: writer_hr_036
Domain: libreoffice_writer

Creates a document with just the company letterhead 'Pinnacle Corp' as the
starting point. The agent must build the full warning letter template.
"""

import os
import shlex
import subprocess
import time

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_036'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_initial():
    doc = Document()

    # Set default page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Company letterhead - Pinnacle Corp
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run('Pinnacle Corp')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)  # Dark navy blue

    # Tagline under company name
    tagline = doc.add_paragraph()
    tagline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tag_run = tagline.add_run('Excellence in Innovation')
    tag_run.italic = True
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Horizontal line separator (thin paragraph border)
    separator = doc.add_paragraph()
    separator.paragraph_format.space_before = Pt(6)
    separator.paragraph_format.space_after = Pt(6)
    pPr = separator._element.get_or_add_pPr()
    pBdr = pPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr', {})
    bottom = pBdr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom', {
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'single',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz': '6',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space': '1',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': '1F3A5F',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # Launch LibreOffice Writer with the document
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
