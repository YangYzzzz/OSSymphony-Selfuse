"""
Reward Script: Written warning letter template with merge fields
Task ID: writer_hr_036
Domain: libreoffice_writer
Scoring:
  Component 1: All 7 required placeholders present (0.45)
  Component 2: "Written Warning" subject line (0.15)
  Component 3: Professional greeting with [Employee Name] (0.10)
  Component 4: Signature block present (0.15)
  Component 5: Document has substantial letter content beyond initial (0.15)
"""

import os
import re

WORKDIR = '/home/user'
TASK_ID = 'writer_hr_036'


def persist_app_state(domain):
    """Save any unsaved LibreOffice changes."""
    import time
    os.environ["DISPLAY"] = ":0"
    if domain in {"libreoffice_calc", "libreoffice_writer", "libreoffice_impress"}:
        try:
            import pyautogui
            pyautogui.hotkey("ctrl", "s")
            time.sleep(0.8)
            print(f"PERSIST: ctrl+s sent for {domain}")
        except Exception as e:
            print(f"PERSIST_WARN: save hook failed: {e}")


def verify_task(file_path):
    """
    Verify task completion with progressive scoring.
    Returns: float between 0.0 and 1.0
    """
    total_score = 0.0

    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        print(f"CRITICAL: Cannot load file {file_path}: {e}")
        print("REWARD: 0.0")
        return 0.0

    # Collect full text of document
    full_text = "\n".join(para.text for para in doc.paragraphs)
    full_text_lower = full_text.lower()

    # Component 1: All 7 required placeholders present (0.45 points)
    # Each placeholder is worth ~0.064 points
    try:
        placeholders = [
            "[Employee Name]",
            "[Employee ID]",
            "[Department]",
            "[Date of Incident]",
            "[Description of Violation]",
            "[Previous Warnings]",
            "[Corrective Action Required]",
        ]
        found_count = 0
        for ph in placeholders:
            if ph in full_text:
                found_count += 1
                print(f"  FOUND: {ph}")
            else:
                print(f"  MISSING: {ph}")

        if found_count == 7:
            print(f"PASS: Component 1 — All 7 placeholders present (0.45 pts)")
            total_score += 0.45
        elif found_count > 0:
            partial = round(0.45 * found_count / 7, 3)
            print(f"PARTIAL: Component 1 — {found_count}/7 placeholders found ({partial} pts)")
            total_score += partial
        else:
            print(f"FAIL: Component 1 — No placeholders found")
    except Exception as e:
        print(f"ERROR: Component 1 — {e}")

    # Component 2: "Written Warning" subject line present (0.15 points)
    # The golden doc has "SUBJECT: Written Warning" - check for "written warning" in text
    try:
        has_written_warning = False
        for para in doc.paragraphs:
            txt = para.text.strip()
            if "written warning" in txt.lower() and ("subject" in txt.lower() or txt.upper() == txt or "warning" in txt.lower()):
                has_written_warning = True
                break
        if has_written_warning:
            print(f"PASS: Component 2 — 'Written Warning' subject line found (0.15 pts)")
            total_score += 0.15
        else:
            print(f"FAIL: Component 2 — No 'Written Warning' subject/title found")
    except Exception as e:
        print(f"ERROR: Component 2 — {e}")

    # Component 3: Professional greeting addressing [Employee Name] (0.10 points)
    # e.g., "Dear [Employee Name]," — must reference the placeholder in a greeting
    try:
        has_greeting = False
        for para in doc.paragraphs:
            txt = para.text.strip().lower()
            if ("dear" in txt or "to:" in txt) and "[employee name]" in txt.lower():
                # Check against actual text (case-sensitive placeholder)
                if "[Employee Name]" in para.text:
                    has_greeting = True
                    break
        if has_greeting:
            print(f"PASS: Component 3 — Professional greeting with [Employee Name] (0.10 pts)")
            total_score += 0.10
        else:
            print(f"FAIL: Component 3 — No greeting addressing [Employee Name]")
    except Exception as e:
        print(f"ERROR: Component 3 — {e}")

    # Component 4: Signature block present (0.15 points)
    # Must have signature lines (Employee Signature, Manager Signature, or similar)
    try:
        signature_keywords = ["signature", "signed", "acknowledged"]
        sig_count = 0
        for para in doc.paragraphs:
            txt = para.text.strip().lower()
            if any(kw in txt for kw in signature_keywords):
                sig_count += 1

        if sig_count >= 2:
            print(f"PASS: Component 4 — Signature block found ({sig_count} sig lines) (0.15 pts)")
            total_score += 0.15
        elif sig_count == 1:
            print(f"PARTIAL: Component 4 — Only 1 signature line found (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 4 — No signature block found")
    except Exception as e:
        print(f"ERROR: Component 4 — {e}")

    # Component 5: Document has substantial letter content beyond initial state (0.15 points)
    # Initial doc has 3 paragraphs. A proper warning letter should have many more.
    # We check that there are at least 15 non-empty paragraphs (golden has 36 total)
    try:
        non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
        if non_empty >= 12:
            print(f"PASS: Component 5 — Substantial content: {non_empty} non-empty paragraphs (0.15 pts)")
            total_score += 0.15
        elif non_empty >= 6:
            print(f"PARTIAL: Component 5 — Some content: {non_empty} non-empty paragraphs (0.08 pts)")
            total_score += 0.08
        else:
            print(f"FAIL: Component 5 — Only {non_empty} non-empty paragraphs, not enough for a letter")
    except Exception as e:
        print(f"ERROR: Component 5 — {e}")

    final_score = min(round(total_score, 2), 1.0)
    print(f"\nScore: {total_score}/1.0")
    print(f"REWARD: {final_score}")
    return final_score


# Entry point
persist_app_state("libreoffice_writer")

file_path = f'{WORKDIR}/{TASK_ID}.docx'
if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    print("REWARD: 0.0")
else:
    verify_task(file_path)
