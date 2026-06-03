"""
Initial Setup: Create a Writer document with five chapters as separate sections,
with footnotes numbered continuously from 1 to 15 across all chapters.
Task ID: writer_bs_030
Domain: libreoffice_writer
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn

WORKDIR = '/home/user'
TASK_ID = 'writer_bs_030'
OUTPUT = f'{WORKDIR}/{TASK_ID}.docx'

def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def add_footnote(paragraph, footnote_text):
    """Add a footnote to a paragraph using raw XML manipulation."""
    run = paragraph.add_run()
    # Create footnote reference in document body
    rPr = run._element.makeelement(qn('w:rPr'), {})
    rStyle = rPr.makeelement(qn('w:rStyle'), {qn('w:val'): 'FootnoteReference'})
    rPr.append(rStyle)
    superscript = rPr.makeelement(qn('w:vertAlign'), {qn('w:val'): 'superscript'})
    rPr.append(superscript)
    run._element.insert(0, rPr)

    footnoteRef = run._element.makeelement(qn('w:footnoteReference'), {})
    # We'll set the ID after creating the footnote element
    run._element.append(footnoteRef)

    # Access the footnotes part
    doc = paragraph.part
    # Get or create footnotes part
    if not hasattr(doc, '_footnotes_part_cache'):
        # Find footnotes relationship
        for rel in doc.rels.values():
            if 'footnotes' in rel.reltype:
                doc._footnotes_part_cache = rel.target_part
                break
        else:
            doc._footnotes_part_cache = None

    if doc._footnotes_part_cache is None:
        # Need to create footnotes part from scratch
        from docx.opc.part import Part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import lxml.etree as etree

        footnotes_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
        partname = '/word/footnotes.xml'
        from docx.opc.packuri import PackURI
        footnotes_part = Part(
            PackURI(partname),
            content_type,
            footnotes_xml.encode('utf-8'),
            doc.package,
        )
        doc.relate_to(footnotes_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
        doc._footnotes_part_cache = footnotes_part

    fn_part = doc._footnotes_part_cache
    import lxml.etree as etree
    fn_root = etree.fromstring(fn_part.blob)

    # Find the next footnote ID
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    existing = fn_root.findall('.//w:footnote', nsmap)
    max_id = 0
    for fn in existing:
        fid = fn.get(qn('w:id'))
        if fid and int(fid) > max_id:
            max_id = int(fid)
    new_id = max_id + 1

    # Create footnote element
    footnote_el = etree.SubElement(fn_root, qn('w:footnote'))
    footnote_el.set(qn('w:id'), str(new_id))

    # Add paragraph inside footnote
    fn_para = etree.SubElement(footnote_el, qn('w:p'))

    # Footnote reference run
    fn_run_ref = etree.SubElement(fn_para, qn('w:r'))
    fn_rPr = etree.SubElement(fn_run_ref, qn('w:rPr'))
    fn_rStyle = etree.SubElement(fn_rPr, qn('w:rStyle'))
    fn_rStyle.set(qn('w:val'), 'FootnoteReference')
    fn_ref = etree.SubElement(fn_run_ref, qn('w:footnoteRef'))

    # Footnote text run
    fn_run_text = etree.SubElement(fn_para, qn('w:r'))
    fn_t = etree.SubElement(fn_run_text, qn('w:t'))
    fn_t.set(qn('xml:space'), 'preserve')
    fn_t.text = ' ' + footnote_text

    # Save back
    fn_part._blob = etree.tostring(fn_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Set reference ID in body
    footnoteRef.set(qn('w:id'), str(new_id))

    return new_id


def create_initial():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ===== Chapter 1: Introduction to Climate Science =====
    heading = doc.add_heading('Chapter 1: Introduction to Climate Science', level=1)

    p1 = doc.add_paragraph(
        'The study of climate science has evolved significantly over the past century. '
        'Early researchers relied on limited observational data from weather stations scattered '
        'across populated regions, leaving vast areas of the globe unmonitored.'
    )
    add_footnote(p1, 'See Arrhenius, S. (1896). On the Influence of Carbonic Acid in the Air upon the Temperature of the Ground.')

    p2 = doc.add_paragraph(
        'Modern climate models incorporate data from satellites, ocean buoys, and atmospheric '
        'sensors to create comprehensive simulations of the Earth\'s climate system. These models '
        'have become increasingly sophisticated, capable of resolving features at scales of '
        'tens of kilometers.'
    )
    add_footnote(p2, 'IPCC Fifth Assessment Report, Working Group I, Chapter 9: Evaluation of Climate Models (2013).')

    p3 = doc.add_paragraph(
        'The Intergovernmental Panel on Climate Change has published six assessment reports since '
        '1990, each building upon the accumulated evidence from thousands of peer-reviewed studies. '
        'The scientific consensus on anthropogenic climate change has strengthened with each successive report.'
    )
    add_footnote(p3, 'Oreskes, N. (2004). The Scientific Consensus on Climate Change. Science, 306(5702), 1686.')

    # ===== Chapter 2: Ocean Circulation Patterns =====
    # New section
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    heading2 = doc.add_heading('Chapter 2: Ocean Circulation Patterns', level=1)

    p4 = doc.add_paragraph(
        'The Atlantic Meridional Overturning Circulation (AMOC) plays a critical role in '
        'regulating global climate. This vast system of ocean currents transports warm surface '
        'water from the tropics northward, where it cools, becomes denser, and sinks to great depths.'
    )
    add_footnote(p4, 'Rahmstorf, S. (2002). Ocean circulation and climate during the past 120,000 years. Nature, 419, 207-214.')

    p5 = doc.add_paragraph(
        'Recent observations suggest that the AMOC has weakened by approximately 15% since the '
        'mid-twentieth century. This slowdown has been linked to increased freshwater input from '
        'melting Greenland ice sheets, which reduces the salinity and density of surface waters '
        'in the North Atlantic.'
    )
    add_footnote(p5, 'Caesar, L. et al. (2018). Observed fingerprint of a weakening Atlantic Ocean overturning circulation. Nature, 556, 191-196.')

    p6 = doc.add_paragraph(
        'The thermohaline circulation extends beyond the Atlantic, forming part of a global '
        'conveyor belt that connects all major ocean basins. Changes in any segment of this system '
        'can propagate effects across the entire network over timescales of decades to centuries.'
    )
    add_footnote(p6, 'Broecker, W.S. (1991). The Great Ocean Conveyor. Oceanography, 4(2), 79-89.')

    p7 = doc.add_paragraph(
        'Pacific Decadal Oscillation (PDO) patterns have shown a correlation with regional '
        'precipitation and temperature anomalies across western North America and eastern Asia, '
        'affecting agricultural productivity in these regions.'
    )
    add_footnote(p7, 'Mantua, N.J. et al. (1997). A Pacific interdecadal climate oscillation. Bulletin of the AMS, 78, 1069-1079.')

    # ===== Chapter 3: Atmospheric Composition =====
    new_section2 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    heading3 = doc.add_heading('Chapter 3: Atmospheric Composition', level=1)

    p8 = doc.add_paragraph(
        'Carbon dioxide concentrations in the atmosphere have risen from approximately 280 parts '
        'per million (ppm) in pre-industrial times to over 420 ppm as of 2024. This increase of '
        'roughly 50% represents the fastest rate of change in atmospheric CO2 in at least 800,000 years '
        'of ice core records.'
    )
    add_footnote(p8, 'Keeling, C.D. et al. (1976). Atmospheric carbon dioxide variations at Mauna Loa Observatory, Hawaii. Tellus, 28(6), 538-551.')

    p9 = doc.add_paragraph(
        'Methane, though present in much lower concentrations than CO2, has a global warming '
        'potential approximately 80 times greater over a 20-year period. Agricultural activities, '
        'particularly rice cultivation and livestock farming, contribute significantly to anthropogenic '
        'methane emissions.'
    )
    add_footnote(p9, 'Saunois, M. et al. (2020). The Global Methane Budget 2000-2017. Earth System Science Data, 12, 1561-1623.')

    # ===== Chapter 4: Polar Ice Dynamics =====
    new_section3 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    heading4 = doc.add_heading('Chapter 4: Polar Ice Dynamics', level=1)

    p10 = doc.add_paragraph(
        'The Greenland Ice Sheet contains enough water to raise global sea levels by approximately '
        '7.2 meters if fully melted. Satellite measurements from the GRACE mission have documented '
        'an accelerating rate of mass loss, averaging 286 billion tonnes per year between 2010 and 2019.'
    )
    add_footnote(p10, 'Mouginot, J. et al. (2019). Forty-six years of Greenland Ice Sheet mass balance. PNAS, 116(19), 9239-9244.')

    p11 = doc.add_paragraph(
        'Antarctic ice dynamics are more complex than those of Greenland, involving both marine-based '
        'and land-based ice sheets. The West Antarctic Ice Sheet (WAIS) is particularly vulnerable '
        'due to its bedrock lying below sea level, making it susceptible to marine ice sheet instability.'
    )
    add_footnote(p11, 'Joughin, I. et al. (2014). Marine Ice Sheet Collapse Potentially Under Way for the Thwaites Glacier Basin. Science, 344(6185), 735-738.')

    p12 = doc.add_paragraph(
        'Arctic sea ice extent has declined by approximately 13% per decade since satellite records '
        'began in 1979. The September minimum extent, which represents the annual low point, has '
        'decreased even more dramatically, with some projections indicating ice-free Arctic summers '
        'by the middle of this century.'
    )
    add_footnote(p12, 'Stroeve, J. et al. (2012). Trends in Arctic sea ice extent from CMIP5, CMIP3 and observations. Geophysical Research Letters, 39(16).')

    # ===== Chapter 5: Mitigation Strategies =====
    new_section4 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    heading5 = doc.add_heading('Chapter 5: Mitigation Strategies', level=1)

    p13 = doc.add_paragraph(
        'Renewable energy deployment has accelerated dramatically in the past decade, with global '
        'solar photovoltaic capacity increasing from 40 GW in 2010 to over 1,200 GW by the end of '
        '2023. The levelized cost of solar energy has fallen by approximately 90% during this period, '
        'making it the cheapest source of new electricity generation in most markets.'
    )
    add_footnote(p13, 'IRENA (2024). Renewable Power Generation Costs in 2023. International Renewable Energy Agency, Abu Dhabi.')

    p14 = doc.add_paragraph(
        'Carbon capture and storage (CCS) technologies remain a contentious element of climate '
        'mitigation portfolios. While proponents argue that CCS is necessary to address emissions '
        'from hard-to-abate industrial sectors, critics point to the technology\'s high costs, '
        'energy penalties, and limited deployment track record.'
    )
    add_footnote(p14, 'Global CCS Institute (2023). Global Status of CCS 2023. Melbourne, Australia.')

    p15 = doc.add_paragraph(
        'Nature-based solutions, including reforestation, peatland restoration, and improved soil '
        'management, offer significant co-benefits beyond carbon sequestration. These approaches can '
        'enhance biodiversity, improve water quality, and support local livelihoods, though their '
        'carbon permanence remains a subject of ongoing research.'
    )
    add_footnote(p15, 'Griscom, B.W. et al. (2017). Natural climate solutions. PNAS, 114(44), 11645-11650.')

    doc.save(OUTPUT)
    print(f'Initial file created: {OUTPUT}')

    # GUI-ready startup
    launch_gui(f'libreoffice --writer "{OUTPUT}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer with DISPLAY=:0')


create_initial()
