"""
Initial Setup: Color correction workflow — docx brief + food photo
Task ID: osworld_multi_apps_writer_gimp_072
Domain: libreoffice_writer + gimp (multi-app)
"""

import os
import shlex
import subprocess
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw
import numpy as np

WORKDIR = '/home/user'
TASK_ID = 'osworld_multi_apps_writer_gimp_072'
DOCX_PATH = f'{WORKDIR}/color_correction_brief.docx'
PHOTO_PATH = f'{WORKDIR}/Desktop/food_photo.jpg'


def launch_gui(command: str, delay_sec: float = 1.0):
    """Launch GUI app on VM display without blocking script exit."""
    env = os.environ.copy()
    env["DISPLAY"] = ":0"
    subprocess.Popen(
        shlex.split(command),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        env=env,
    )
    time.sleep(delay_sec)


def create_docx():
    """Create the color_correction_brief.docx with detailed correction steps."""
    doc = Document()

    # Title
    title = doc.add_heading('Color Correction Brief — Food Photography', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro paragraph
    intro = doc.add_paragraph(
        'This brief outlines the specific color grading adjustments to be applied to the '
        'food photograph (food_photo.jpg) using GIMP. Follow the steps below in order to '
        'achieve the desired warm, vibrant look suitable for publication.'
    )

    doc.add_paragraph('')  # spacer

    # Step 1: Selective Color
    doc.add_heading('Step 1: Selective Color', level=1)
    p1 = doc.add_paragraph(
        'Navigate to Colors > Curves (or use the Selective Color tool). '
        'In the Reds channel range, apply the following adjustments:'
    )
    step1_items = [
        'Red channel: +15 (boost red in highlights)',
        'Cyan channel: -10 (reduce cyan in the reds range)',
    ]
    for item in step1_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')

    # Step 2: Color Balance
    doc.add_heading('Step 2: Color Balance — Shadows', level=1)
    p2 = doc.add_paragraph(
        'Open Colors > Color Balance. Select the Shadows tone range and apply:'
    )
    step2_items = [
        'Cyan/Red slider: -8 (shift slightly away from cyan toward neutral in shadows)',
    ]
    for item in step2_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')

    # Step 3: Hue/Saturation Warmth
    doc.add_heading('Step 3: Hue/Saturation — Overall Warmth', level=1)
    p3 = doc.add_paragraph(
        'Open Colors > Hue-Saturation. With All channels selected, apply a subtle '
        'warmth enhancement via hue rotation:'
    )
    step3_items = [
        'Hue rotation: +8 degrees (shifts color temperature toward warm tones)',
        'This simulates a warm filter effect across the entire image',
    ]
    for item in step3_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')

    # Final note
    note = doc.add_paragraph(
        'Save the result as food_corrected.jpg on the Desktop. '
        'These adjustments are designed to enhance the natural warmth of the food and '
        'reduce the cool/cyan cast in shadows that is common in studio photography.'
    )
    run = note.runs[0] if note.runs else note.add_run()
    # Note paragraph is plain, add run with italic note
    doc.add_paragraph('')
    note2 = doc.add_paragraph()
    run2 = note2.add_run('Note: Apply these steps in sequence for the best result.')
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    doc.save(DOCX_PATH)
    print(f'Created: {DOCX_PATH}')


def create_food_photo():
    """Create a realistic-looking food photo (colorful dish scene)."""
    os.makedirs(os.path.dirname(PHOTO_PATH), exist_ok=True)

    # Create a 800x600 RGB image representing a food scene
    img = Image.new('RGB', (800, 600), color=(245, 240, 220))
    draw = ImageDraw.Draw(img)

    # Background: warm table surface
    for y in range(600):
        r = int(220 + (y / 600) * 20)
        g = int(195 + (y / 600) * 15)
        b = int(160 + (y / 600) * 10)
        draw.line([(0, y), (800, y)], fill=(min(255, r), min(255, g), min(255, b)))

    # Plate (white/light gray circular shape)
    draw.ellipse([150, 100, 650, 500], fill=(248, 245, 240), outline=(200, 195, 188), width=3)

    # Main dish — pasta-like item in center (warm orangey-red tones)
    draw.ellipse([220, 170, 580, 430], fill=(210, 120, 60))

    # Some pasta strands (curved lines)
    for i in range(8):
        x0 = 230 + i * 40
        y0 = 190 + (i % 3) * 20
        x1 = x0 + 20
        y1 = y0 + 60
        draw.arc([x0, y0, x1, y1], start=0, end=180, fill=(185, 90, 40), width=3)

    # Sauce / red tomato areas
    draw.ellipse([270, 200, 370, 280], fill=(195, 50, 35))
    draw.ellipse([390, 230, 500, 310], fill=(190, 45, 30))
    draw.ellipse([310, 290, 420, 380], fill=(200, 55, 40))

    # Green basil leaves
    draw.ellipse([300, 195, 340, 225], fill=(60, 140, 50))
    draw.ellipse([440, 210, 480, 238], fill=(65, 145, 55))
    draw.ellipse([350, 320, 385, 348], fill=(55, 135, 45))

    # Parmesan/cheese sprinkle (off-white flecks)
    for cx, cy in [(260, 240), (340, 260), (410, 250), (470, 280), (320, 340), (450, 350)]:
        draw.ellipse([cx-5, cy-3, cx+5, cy+3], fill=(240, 235, 210))

    # Shadow around plate edge
    draw.arc([148, 98, 652, 502], start=200, end=340, fill=(180, 170, 150), width=4)

    # Small garnish items (cherry tomatoes)
    for cx, cy in [(180, 160), (610, 170), (620, 400), (170, 390)]:
        draw.ellipse([cx-12, cy-12, cx+12, cy+12], fill=(200, 40, 30))
        draw.ellipse([cx-5, cy-10, cx+2, cy-5], fill=(230, 80, 70))  # highlight

    # Slight cyan-ish shadow tint in lower areas (gives the typical studio cast)
    shadow_overlay = Image.new('RGBA', (800, 600), (0, 0, 0, 0))
    sd = ImageDraw.Draw(shadow_overlay)
    for y in range(400, 600):
        alpha = int((y - 400) / 200 * 30)
        sd.line([(0, y), (800, y)], fill=(30, 50, 80, alpha))
    img = img.convert('RGBA')
    img = Image.alpha_composite(img, shadow_overlay)
    img = img.convert('RGB')

    img.save(PHOTO_PATH, 'JPEG', quality=92)
    print(f'Created: {PHOTO_PATH}')


def main():
    create_docx()
    create_food_photo()

    # GUI-ready startup: open the docx in LibreOffice Writer, then open food_photo.jpg in GIMP
    launch_gui(f'libreoffice --writer "{DOCX_PATH}"', delay_sec=3.0)
    launch_gui(f'gimp "{PHOTO_PATH}"', delay_sec=2.0)
    print('GUI_READY: launched LibreOffice Writer and GIMP with DISPLAY=:0')


main()
